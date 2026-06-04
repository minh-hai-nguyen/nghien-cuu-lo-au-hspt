# -*- coding: utf-8 -*-
"""Find the existing PHẢN BIỆN/critique section in VN022 docx."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

d = Document(PATH)
total = len(d.paragraphs)
print(f'Total paragraphs: {total}\n')

# Find PHẢN BIỆN or related headings
keywords = ['PHẢN BIỆN', 'ĐÁNH GIÁ', 'BÌNH LUẬN', 'CRITIQUE', 'điểm mạnh', 'hạn chế', 'GAP']
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if any(k.lower() in t.lower() for k in keywords) and len(t) < 200:
        print(f'  #{i}: [{t[:120]}]')
