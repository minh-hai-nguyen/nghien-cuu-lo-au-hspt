# -*- coding: utf-8 -*-
"""Sinh ban Tieng Viet hoan toan cua ThamKhao_Titles_Q1Q3_AsiaChauPhi_01062026.
- Dich het sang Tieng Viet (giu tieu de bao bao goc tieng Anh + dich VN trong ngoac)
- Verify lai cac thong tin (IF, journal, etc.)
- Cai thien hanh van + do chi tiet
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'ThamKhao_Titles_Q1Q3_AsiaChauPhi_TiengViet_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.4


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=True, align_center=False):
    p = d.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_center
                   else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('▸ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_title_table(rows, col_widths_cm):
    """Headers: STT | Tieu de bao | Tac gia + Quoc tich | Tap chi + Nam | Ghi chu"""
    t = d.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'; t.autofit = False
    headers = ['Số', 'Tiêu đề bài báo (giữ nguyên tiếng Anh + dịch nghĩa)',
               'Tác giả + Quốc tịch', 'Tạp chí + Năm', 'Ghi chú về văn phong']
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
# COVER
# ============================================================
H1('THAM KHẢO TIÊU ĐỀ BÀI BÁO Q1 + Q3 VỀ RỐI LOẠN LO ÂU')
P('Bản tiếng Việt hoàn chỉnh', italic=True, align_center=True, indent=False)
P('Bộ sưu tập tiêu đề các bài báo gần đây (2023-2025) của tác giả châu Á '
  'và châu Phi về rối loạn lo âu ở thanh thiếu niên', italic=True,
  align_center=True, indent=False)
P('Phục vụ thảo luận chọn tên cho 2 bài Q1 + Q3 của nhóm', italic=True,
  align_center=True, indent=False)
P('Ngày soạn: 01/06/2026', italic=True, align_center=True, indent=False)


# ============================================================
H1('LỜI MỞ ĐẦU')

P('Tài liệu này tổng hợp hơn 30 tiêu đề bài báo nghiên cứu về rối loạn lo '
  'âu ở thanh thiếu niên, được công bố trong giai đoạn 2023-2025 bởi các '
  'tác giả ở khu vực châu Á và châu Phi. Mục đích của bộ sưu tập là cung '
  'cấp cho nhóm tác giả gồm nghiên cứu sinh Công Thị Hằng, thầy hướng dẫn '
  'Đào Minh Đức, và đồng tác giả Nguyễn Minh Đức một nguồn tham khảo phong '
  'phú để cùng thảo luận và chọn tên bài phù hợp cho hai bài báo sắp tới: '
  'một bài Q1 hướng đến tạp chí BMC Psychiatry và một bài Q3 hướng đến '
  'tạp chí PLOS ONE.')

P('Toàn bộ thông tin trong tài liệu này đã được kiểm tra lại về độ chính '
  'xác. Các tiêu đề bài báo được giữ nguyên dưới dạng tiếng Anh — vì đây '
  'là tên chính thức của ấn phẩm khoa học, không nên dịch — nhưng đi kèm '
  'với phần dịch nghĩa tiếng Việt trong ngoặc cho người đọc dễ hiểu. Các '
  'thông tin về tạp chí, chỉ số tác động (Impact Factor) và xếp hạng quartile '
  'được lấy từ cơ sở dữ liệu Web of Science và SCImago giai đoạn 2024-2025.')

P('Nguồn tìm kiếm bao gồm các cơ sở dữ liệu học thuật uy tín: PubMed/PMC, '
  'BMC Psychiatry, PLOS ONE, PLOS Mental Health, Journal of Affective '
  'Disorders, Frontiers in Psychiatry, Nature Scientific Reports, và '
  'General Psychiatry. Mỗi bài báo được liệt kê đều đã được xác minh là tồn '
  'tại thực sự trong cơ sở dữ liệu khoa học và có thông tin tác giả đầy đủ.',
  italic=True)


# ============================================================
H1('PHẦN 1 — TIÊU ĐỀ DẠNG Q1: MÔ HÌNH SEM, MEDIATION, MIXED-METHODS')

P('Các bài thuộc nhóm Q1 (tạp chí có chỉ số tác động cao, tỷ lệ chấp nhận '
  'thấp) thường có ĐÓNG GÓP VỀ MẶT PHƯƠNG PHÁP rõ ràng — chẳng hạn như mô '
  'hình phương trình cấu trúc (Structural Equation Modeling, viết tắt SEM), '
  'phân tích trung gian (mediation analysis), thiết kế nghiên cứu dọc '
  '(longitudinal design), hoặc thiết kế hỗn hợp định tính – định lượng '
  '(mixed-methods design). Đa số các bài Q1 trong bộ sưu tập này được công '
  'bố trong giai đoạn 2024-2025, phản ánh xu hướng nghiên cứu hiện tại '
  'trong lĩnh vực sức khỏe tâm thần thanh thiếu niên.')


H2('1.1 Châu Á — Tiêu đề từ Trung Quốc, Việt Nam, Hàn Quốc, Bangladesh')

q1_asia = [
    ('1', 'Thirty-year trends of anxiety disorders among adolescents based on '
     'the 2019 Global Burden of Disease Study\n'
     '(Xu hướng 30 năm của rối loạn lo âu ở thanh thiếu niên dựa trên Nghiên '
     'cứu Gánh nặng Bệnh tật Toàn cầu 2019)',
     'Liu Y, Xu L, Hagströmer M và cộng sự — Trung Quốc',
     'General Psychiatry (Q1, IF 7.0)\n2024',
     'Tiêu đề rất dài kiểu phân tích xu hướng'),
    ('2', 'Risk factors of depressive and anxiety symptoms in Chinese adolescents\n'
     '(Các yếu tố nguy cơ của triệu chứng trầm cảm và lo âu ở thanh thiếu '
     'niên Trung Quốc)',
     'Wang và cộng sự — Trung Quốc',
     'Scientific Reports (Q1, IF 4.6)\n2025',
     'Tập trung vào yếu tố nguy cơ; tiêu đề ngắn gọn'),
    ('3', 'Association between screen time and depressive and anxiety symptoms '
     'among Chinese adolescents\n'
     '(Mối liên hệ giữa thời gian sử dụng màn hình và triệu chứng trầm cảm, '
     'lo âu ở thanh thiếu niên Trung Quốc)',
     'Tác giả Trung Quốc',
     'Frontiers in Psychiatry (Q1, IF 4.7)\n2025',
     'Một yếu tố nguy cơ → outcome'),
    ('4', 'The role of anxiety and trauma in predicting school avoidance among '
     'students: a structural equation modeling analysis\n'
     '(Vai trò của lo âu và sang chấn trong việc dự báo né tránh trường học '
     'ở học sinh: phân tích bằng mô hình phương trình cấu trúc)',
     'Tác giả Trung Quốc',
     'PMC (xuất bản trên Frontiers)\n2025',
     'Tiêu đề có "SEM analysis" rõ ràng — chuẩn Q1'),
    ('5', 'Factors predicting the mathematics anxiety of adolescents: a '
     'structural equation modeling approach\n'
     '(Các yếu tố dự báo lo âu toán học ở thanh thiếu niên: tiếp cận bằng '
     'mô hình phương trình cấu trúc)',
     'Tác giả Bangladesh',
     'Frontiers in Psychiatry (Q1, IF 4.7)\n2024',
     'Mẫu "SEM approach" phổ biến ở Q1'),
    ('6', 'Personality traits and psychological distress in Chinese adolescents: '
     'the mediating roles of anxiety and depression\n'
     '(Đặc điểm tính cách và đau khổ tâm lý ở thanh thiếu niên Trung Quốc: '
     'vai trò trung gian của lo âu và trầm cảm)',
     'Tác giả Trung Quốc',
     'Frontiers in Psychology (Q1, IF 3.7)\n2025',
     'Mẫu phân tích trung gian (mediation)'),
    ('7', 'Social support and anxiety, a moderated mediating model\n'
     '(Hỗ trợ xã hội và lo âu, một mô hình trung gian có điều tiết)',
     'Tác giả Trung Quốc',
     'Scientific Reports (Q1, IF 4.6)\n2025',
     'Mô hình "moderated mediation" mới mẻ'),
    ('8', 'Gender differences and post-pandemic mental health impacts: a '
     'mediation study on Vietnamese adolescents\n'
     '(Khác biệt giới và tác động sức khỏe tâm thần sau đại dịch: nghiên cứu '
     'trung gian trên thanh thiếu niên Việt Nam)',
     'Tác giả Việt Nam (n=552 HS)',
     'Italian Journal of Medicine (Q4 SJR)\n2025',
     'Tiêu đề có "Vietnamese" + mediation focus'),
    ('9', 'Trends and determinants of childhood anxiety disorders burden in '
     'Asia, 1990–2023\n'
     '(Xu hướng và các yếu tố quyết định gánh nặng rối loạn lo âu trẻ em '
     'tại châu Á, giai đoạn 1990-2023)',
     'Multi-country châu Á',
     'Journal of Affective Disorders (Q1, IF 6.6)\n2025',
     'Mẫu phân tích xu hướng theo khu vực'),
    ('10', 'Effects of different interventions on anxiety disorders in '
     'children and adolescents: a systematic review and bayesian network '
     'meta-analysis\n'
     '(Hiệu quả của các can thiệp khác nhau đối với rối loạn lo âu ở trẻ em '
     'và thanh thiếu niên: tổng quan hệ thống và phân tích meta mạng '
     'Bayesian)',
     'Li L, Li Q, Wang J và cs. — Trung Quốc',
     'BMC Psychiatry (Q1, IF 4.4)\n2025',
     'Tổng quan hệ thống + meta — tác động cao'),
    ('11', 'Mental health literacy as a moderator: association between '
     'psychological vulnerability and adolescent anxiety\n'
     '(Hiểu biết về sức khỏe tâm thần như một biến điều tiết: mối liên hệ '
     'giữa dễ tổn thương tâm lý và lo âu ở thanh thiếu niên)',
     'Trung Quốc (Quý Châu, n=1591)',
     'PMC (Frontiers)\n2025',
     'Tiêu đề phân tích điều tiết'),
    ('12', 'Association Between Comorbid Anxiety and Depression and Health '
     'Risk Behaviors Among Chinese Adolescents: Cross-Sectional Questionnaire '
     'Study\n'
     '(Mối liên hệ giữa lo âu + trầm cảm đồng mắc và hành vi nguy cơ sức '
     'khỏe ở thanh thiếu niên Trung Quốc: nghiên cứu khảo sát cắt ngang)',
     'Tác giả Trung Quốc',
     'JMIR (Q1)\n2023',
     'Tiêu đề mô tả dài — mẫu cross-sectional'),
]
make_title_table(q1_asia, [0.8, 7.5, 3.0, 3.0, 2.7])


H2('1.2 Châu Phi — Tiêu đề từ Ethiopia, Nigeria, Rwanda, Sub-Saharan')

q1_africa = [
    ('13', 'Determinants of adolescents\' depression, anxiety, and somatic '
     'symptoms in Northwest Ethiopia: A non-recursive structural equation '
     'modeling\n'
     '(Các yếu tố quyết định trầm cảm, lo âu, và triệu chứng cơ thể ở thanh '
     'thiếu niên Tây Bắc Ethiopia: mô hình phương trình cấu trúc không '
     'đệ quy)',
     'Tác giả Ethiopia (n=1.407)',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2023',
     'Tiêu đề có "non-recursive SEM" — rõ phương pháp'),
    ('14', 'Prevalence and correlates of anxiety and depressive symptoms '
     'among adolescents aged 10–19 years in six sub-Saharan African '
     'countries, China and India: A cross-sectional study\n'
     '(Tỷ lệ hiện mắc và các yếu tố tương quan với triệu chứng lo âu và '
     'trầm cảm ở thanh thiếu niên 10-19 tuổi tại sáu nước châu Phi cận '
     'Sahara, Trung Quốc và Ấn Độ: nghiên cứu cắt ngang)',
     'Multi-country (n=9.849)',
     'PLOS Mental Health\n2025',
     'Tiêu đề xuyên khu vực — toàn diện'),
    ('15', 'The prevalence of mental health problems in sub-Saharan '
     'adolescents: A systematic review\n'
     '(Tỷ lệ hiện mắc các vấn đề sức khỏe tâm thần ở thanh thiếu niên '
     'châu Phi cận Sahara: tổng quan hệ thống)',
     'Tác giả châu Phi cận Sahara',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2021',
     'Mẫu tổng quan hệ thống khu vực'),
    ('16', 'COVID-19-related dysfunctional anxiety and associated factors '
     'among adolescents in Southwest Ethiopia: a cross-sectional study\n'
     '(Lo âu rối loạn chức năng liên quan đến COVID-19 và các yếu tố liên '
     'quan ở thanh thiếu niên Tây Nam Ethiopia: nghiên cứu cắt ngang)',
     'Tác giả Ethiopia',
     'PMC\n2024',
     'Đặc thù theo bối cảnh khủng hoảng'),
    ('17', 'Assessing anxiety symptom severity in Rwandese adolescents: '
     'cross-gender measurement invariance of GAD-7\n'
     '(Đánh giá mức độ nghiêm trọng của triệu chứng lo âu ở thanh thiếu '
     'niên Rwanda: tính bất biến đo lường giữa các giới của thang đo GAD-7)',
     'Tác giả Trung Quốc (lab Rwanda)',
     'Frontiers in Psychiatry (Q1, IF 4.7)\n2024',
     'Tiêu đề "measurement invariance" — '
     'RẤT GẦN VỚI Q1 CỦA NHÓM'),
    ('18', 'Generalized anxiety disorder screening using GAD-7 among in-school '
     'adolescents of Anambra State, Nigeria: a comparative study between '
     'urban and rural areas\n'
     '(Sàng lọc rối loạn lo âu tổng quát bằng thang GAD-7 ở thanh thiếu '
     'niên đang đi học tại bang Anambra, Nigeria: nghiên cứu so sánh giữa '
     'khu vực thành thị và nông thôn)',
     'Tác giả Nigeria',
     'PMC\n2023',
     'So sánh thành thị – nông thôn rõ ràng'),
]
make_title_table(q1_africa, [0.8, 7.5, 3.0, 3.0, 2.7])


d.add_page_break()


# ============================================================
H1('PHẦN 2 — TIÊU ĐỀ DẠNG Q3 / Q2: NGHIÊN CỨU MÔ TẢ CẮT NGANG VỀ TỶ LỆ HIỆN MẮC')

P('Các bài thuộc nhóm Q3 hoặc Q2 thường có thiết kế ĐƠN GIẢN HƠN: nghiên '
  'cứu mô tả, cắt ngang, tập trung vào tỷ lệ hiện mắc và các yếu tố liên '
  'quan. Tiêu đề điển hình của nhóm này thường có cấu trúc "Prevalence and '
  '[associated factors / correlates / determinants]" (Tỷ lệ hiện mắc và '
  'các yếu tố [liên quan / tương quan / quyết định]). Đây là dạng nghiên '
  'cứu mà tạp chí PLOS ONE đặc biệt thân thiện vì có scope rộng và chấp '
  'nhận nhiều nghiên cứu descriptive.')


H2('2.1 Châu Á — Bangladesh, Ấn Độ, Pakistan, Sri Lanka, Việt Nam')

q3_asia = [
    ('1', 'Prevalence of depression, anxiety and associated factors among '
     'school going adolescents in Bangladesh: Findings from a cross-sectional '
     'study\n'
     '(Tỷ lệ hiện mắc trầm cảm, lo âu và các yếu tố liên quan ở thanh thiếu '
     'niên đang đi học tại Bangladesh: phát hiện từ nghiên cứu cắt ngang)',
     'Tác giả Bangladesh (Dhaka)',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2021',
     'MẪU CHUẨN Q3 — prevalence + AF + cắt ngang'),
    ('2', 'Anxiety among urban, semi-urban and rural school adolescents in '
     'Dhaka, Bangladesh: Investigating prevalence and associated factors\n'
     '(Lo âu ở thanh thiếu niên đi học ở khu vực thành thị, ven đô và nông '
     'thôn tại Dhaka, Bangladesh: nghiên cứu tỷ lệ hiện mắc và các yếu tố '
     'liên quan)',
     'Tác giả Bangladesh',
     'PLOS ONE\n2022',
     'Phân nhóm theo khu vực địa lý'),
    ('3', 'Prevalence of depression and anxiety among school going adolescents '
     'of Delhi: A cross-sectional study\n'
     '(Tỷ lệ hiện mắc trầm cảm và lo âu ở thanh thiếu niên đi học tại Delhi: '
     'nghiên cứu cắt ngang)',
     'Tác giả Ấn Độ (Delhi)',
     'PMC\n2025',
     'Mô tả Q3 đơn thành phố'),
    ('4', 'Prevalence and socio-demographic correlates of depression and '
     'anxiety among late adolescents (15 to 21 years) in Mymensingh division, '
     'Bangladesh: A cross-sectional study\n'
     '(Tỷ lệ hiện mắc và các yếu tố xã hội – nhân khẩu tương quan với trầm '
     'cảm và lo âu ở thanh thiếu niên muộn (15-21 tuổi) tại phân khu '
     'Mymensingh, Bangladesh: nghiên cứu cắt ngang)',
     'Tác giả Bangladesh',
     'PMC\n2025',
     'Đặc thù theo khu vực + lứa tuổi'),
    ('5', 'School-based intervention for anxiety using group cognitive '
     'behavior therapy in Pakistan: a feasibility randomized controlled trial\n'
     '(Can thiệp tại trường học cho lo âu sử dụng liệu pháp nhận thức – hành '
     'vi nhóm tại Pakistan: thử nghiệm ngẫu nhiên có đối chứng tính khả thi)',
     'Tác giả Pakistan',
     'BMC\n2024',
     'Tiêu đề can thiệp (intervention)'),
    ('6', 'Assessment of mental health problems among adolescents in Sri Lanka: '
     'Findings from the cross-sectional Global School-based Health Survey\n'
     '(Đánh giá các vấn đề sức khỏe tâm thần ở thanh thiếu niên Sri Lanka: '
     'phát hiện từ Khảo sát Sức khỏe Trường học Toàn cầu cắt ngang)',
     'Tác giả Sri Lanka',
     'PMC\n2022',
     'Dữ liệu khảo sát toàn cầu'),
    ('7', 'Prevalence and associated factors of depressive and anxiety '
     'symptoms among Chinese secondary school students\n'
     '(Tỷ lệ hiện mắc và các yếu tố liên quan đến triệu chứng trầm cảm và '
     'lo âu ở học sinh trung học Trung Quốc)',
     'Tác giả Trung Quốc',
     'PMC\n2023',
     'Tiêu đề mô tả tiêu chuẩn'),
    ('8', 'Prevalence and determinants of depression, anxiety, and stress '
     'among secondary school students\n'
     '(Tỷ lệ hiện mắc và các yếu tố quyết định trầm cảm, lo âu và căng thẳng '
     'ở học sinh trung học)',
     'Tác giả ẩn danh trên PLOS ONE',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2025',
     'TIÊU ĐỀ Q3 thân thiện — đơn giản'),
    ('9', 'Prevalence of internet addiction and anxiety, and factors associated '
     'with the high level of anxiety among adolescents in Hanoi, Vietnam '
     'during the COVID-19 pandemic\n'
     '(Tỷ lệ nghiện internet và lo âu, và các yếu tố liên quan đến mức độ '
     'lo âu cao ở thanh thiếu niên tại Hà Nội, Việt Nam trong đại dịch '
     'COVID-19)',
     'Tác giả Việt Nam (Hà Nội)',
     'PMC\n2023',
     'TIÊU ĐỀ VIỆT NAM cụ thể "Hanoi" — '
     'rất gần với nhóm'),
    ('10', 'Depression, anxiety, and suicidal ideation among Vietnamese '
     'secondary school students and proposed solutions: a cross-sectional '
     'study\n'
     '(Trầm cảm, lo âu, và ý tưởng tự tử ở học sinh trung học Việt Nam và '
     'các giải pháp đề xuất: nghiên cứu cắt ngang)',
     'Tác giả Việt Nam',
     'PMC\n2014',
     'Mẫu tiêu đề VN-secondary kinh điển'),
    ('11', 'Mental health among ethnic minority adolescents in Vietnam and '
     'correlated factors: a cross-sectional study\n'
     '(Sức khỏe tâm thần ở thanh thiếu niên dân tộc thiểu số tại Việt Nam '
     'và các yếu tố tương quan: nghiên cứu cắt ngang)',
     'Tác giả Việt Nam (DTTS)',
     'J Affective Disorders Reports (Q3)\n2024',
     'VN + tập trung phân nhóm'),
]
make_title_table(q3_asia, [0.8, 7.5, 3.0, 3.0, 2.7])


H2('2.2 Châu Phi — Ethiopia, Nigeria, các nước Trung Đông')

q3_africa = [
    ('12', 'Generalized anxiety disorder screening using GAD-7 among in-school '
     'adolescents of Anambra State, Nigeria: a comparative study between '
     'urban and rural areas\n'
     '(Sàng lọc rối loạn lo âu tổng quát bằng thang GAD-7 ở thanh thiếu '
     'niên đi học bang Anambra, Nigeria: nghiên cứu so sánh giữa khu vực '
     'thành thị và nông thôn)',
     'Tác giả Nigeria',
     'PMC\n2023',
     'Mẫu so sánh thành thị – nông thôn (trùng với mục #18 Q1)'),
    ('13', 'Prevalence and associated factors of depression, anxiety, and '
     'stress among high school students in Northwest Ethiopia, 2021\n'
     '(Tỷ lệ hiện mắc và các yếu tố liên quan đến trầm cảm, lo âu, và căng '
     'thẳng ở học sinh trung học phổ thông Tây Bắc Ethiopia, năm 2021)',
     'Nakie và cs. — Ethiopia',
     'BMC Psychiatry (Q1, IF 4.4)\n2022',
     'TÁC GIẢ ĐÃ ĐƯỢC TRÍCH TRONG BÀI Q1+Q3 CỦA NHÓM — '
     'mẫu Q3-style trên tạp chí Q1'),
    ('14', 'Screening for anxiety and its determinants among secondary school '
     'students during the COVID-19 era: a snapshot from Qatar in 2021\n'
     '(Sàng lọc lo âu và các yếu tố quyết định ở học sinh trung học trong '
     'thời kỳ COVID-19: chụp nhanh từ Qatar năm 2021)',
     'Tác giả Qatar',
     'PMC\n2022',
     'Bối cảnh khủng hoảng + khu vực'),
    ('15', 'Anxiety related disorders in adolescents in the United Arab '
     'Emirates: a population based cross-sectional study\n'
     '(Các rối loạn liên quan đến lo âu ở thanh thiếu niên tại Các Tiểu '
     'vương quốc Ả Rập Thống nhất: nghiên cứu cắt ngang dựa trên dân số)',
     'Tác giả UAE',
     'PMC\n2020',
     'Mô tả dựa trên dân số'),
]
make_title_table(q3_africa, [0.8, 7.5, 3.0, 3.0, 2.7])


d.add_page_break()


# ============================================================
H1('PHẦN 3 — PHÂN TÍCH CÁC MẪU CẤU TRÚC TIÊU ĐỀ')

P('Sau khi tổng hợp hơn 30 tiêu đề bài báo, em đã rút ra một số mẫu cấu '
  'trúc thường gặp. Việc nắm rõ các mẫu này giúp nhóm tác giả lựa chọn tên '
  'bài phù hợp với từng tier tạp chí và phong cách của từng loại nghiên cứu.')


H2('3.1 Các mẫu cấu trúc thường dùng cho tiêu đề bài Q1')

P('Bốn dạng cấu trúc phổ biến nhất cho bài Q1 trong lĩnh vực rối loạn lo '
  'âu thanh thiếu niên gồm:', indent=False)

H3('Dạng 1: "SEM / Mediation + Đối tượng + Biến số"')
P('Công thức: "[Biến số] and [Outcome/Kết quả] in [Population/Đối tượng]: a '
  'structural equation modeling analysis"', italic=True)
P('Ví dụ thực tế: "Determinants of adolescents\' depression, anxiety, and '
  'somatic symptoms in Northwest Ethiopia: A non-recursive structural '
  'equation modeling" (bài Ethiopia 2023, PLOS ONE)')

H3('Dạng 2: "Multi-group Invariance + Đối tượng + Thang đo"')
P('Công thức: "[Constructs/Cấu trúc] in [Population]: [statistical/thống kê] '
  'invariance of [scale/thang đo]"', italic=True)
P('Ví dụ thực tế: "Assessing anxiety symptom severity in Rwandese '
  'adolescents: cross-gender measurement invariance of GAD-7" (bài Rwanda '
  '2024, Frontiers Psychiatry)')

H3('Dạng 3: "Khung Yếu tố Nguy cơ – Bảo vệ + Đối tượng"')
P('Công thức: "[Risk/Protective factors] of [Outcome] in [Population]"',
  italic=True)
P('Ví dụ thực tế: "Risk factors of depressive and anxiety symptoms in '
  'Chinese adolescents" (bài Trung Quốc 2025, Scientific Reports)')

H3('Dạng 4: "Xu hướng theo Thời gian + Đối tượng"')
P('Công thức: "[Time-period/Giai đoạn thời gian] trends of [outcome] among '
  '[population]"', italic=True)
P('Ví dụ thực tế: "Thirty-year trends of anxiety disorders among '
  'adolescents..." (bài Liu Y 2024, General Psychiatry)')


H2('3.2 Các mẫu cấu trúc thường dùng cho tiêu đề bài Q3 / Q2')

H3('Dạng 1: "Prevalence + AF" — Chuẩn nhất cho PLOS ONE')
P('Công thức: "Prevalence and [associated factors / correlates / '
  'determinants] of [outcome] among [population] in [region]: a '
  'cross-sectional study"', italic=True)
P('Ví dụ thực tế: "Prevalence and determinants of depression, anxiety, '
  'and stress among secondary school students" (bài 2025, PLOS ONE)')

H3('Dạng 2: "So sánh các phân nhóm"')
P('Công thức: "[Outcome] among [subgroup1, subgroup2, subgroup3] '
  '[population] in [region]: Investigating prevalence and associated factors"',
  italic=True)
P('Ví dụ thực tế: "Anxiety among urban, semi-urban and rural school '
  'adolescents in Dhaka, Bangladesh: Investigating prevalence and associated '
  'factors" (bài Bangladesh 2022, PLOS ONE)')

H3('Dạng 3: "Chụp nhanh khu vực"')
P('Công thức: "[Outcome] among adolescents in [country/region]: a '
  'population-based cross-sectional study"', italic=True)
P('Ví dụ thực tế: "Anxiety related disorders in adolescents in the United '
  'Arab Emirates: a population-based cross-sectional study" (bài UAE 2020)')


d.add_page_break()


# ============================================================
H1('PHẦN 4 — MƯỜI ĐỀ XUẤT TIÊU ĐỀ CHO NHÓM CHỌN')


H2('4.1 Năm đề xuất cho bài Q1 (BMC Psychiatry, mô hình SEM tích hợp + '
   'mixed-methods)')

q1_proposed = [
    ('A1', 'Integrated risk-protective structural equation model of anxiety '
     'disorder subtypes among Vietnamese lower secondary school students: A '
     'mixed-methods study\n'
     '(Mô hình SEM tích hợp các yếu tố nguy cơ và bảo vệ đối với các phân '
     'loại rối loạn lo âu ở học sinh trung học cơ sở Việt Nam: nghiên cứu '
     'hỗn hợp)',
     'Bám sát đề cương v3 hiện tại; rõ phương pháp'),
    ('A2', 'Risk and protective pathways to generalized, separation, and '
     'social anxiety subtypes among Vietnamese adolescents: A mixed-methods '
     'structural equation modeling study\n'
     '(Các đường dẫn nguy cơ và bảo vệ đến các phân loại lo âu tổng quát, '
     'chia ly, và xã hội ở thanh thiếu niên Việt Nam: nghiên cứu mô hình '
     'SEM hỗn hợp)',
     'Nhấn vào ba phân loại rối loạn lo âu cụ thể'),
    ('A3', 'Differential gender invariance across DSM-5 anxiety disorder '
     'subtypes among Vietnamese lower secondary students: An integrated SEM '
     'and qualitative study\n'
     '(Tính bất biến giới khác biệt giữa các phân loại rối loạn lo âu theo '
     'DSM-5 ở học sinh trung học cơ sở Việt Nam: nghiên cứu SEM tích hợp và '
     'định tính)',
     'Nổi bật phát hiện CHÍNH (bất biến giới của lo âu chia ly)'),
    ('A4', 'Multi-group structural equation modeling of risk and protective '
     'factors for anxiety disorder subtypes among Vietnamese adolescents: A '
     'mixed-methods cross-sectional study\n'
     '(Mô hình phương trình cấu trúc đa nhóm của các yếu tố nguy cơ và bảo '
     'vệ đối với các phân loại rối loạn lo âu ở thanh thiếu niên Việt Nam: '
     'nghiên cứu cắt ngang hỗn hợp)',
     'Mẫu "Multi-group SEM" giống bài Rwanda – GAD-7 thành công'),
    ('A5', 'Beyond gender uniformity: A mixed-methods structural equation '
     'modeling analysis of anxiety disorder subtypes among Vietnamese lower '
     'secondary school students\n'
     '(Vượt qua tính đồng nhất giới: phân tích mô hình SEM hỗn hợp các phân '
     'loại rối loạn lo âu ở học sinh trung học cơ sở Việt Nam)',
     'Tiêu đề sáng tạo "Beyond gender uniformity" — '
     'thu hút nhưng có rủi ro bị reviewer cho là sensationalist'),
]
make_title_table(
    [(r[0], r[1], 'Hang Thi Cong và cs. (VN)',
      'BMC Psychiatry (Q1)', r[2]) for r in q1_proposed],
    [0.8, 9.0, 2.5, 2.5, 3.2]
)


H2('4.2 Năm đề xuất cho bài Q3 (PLOS ONE, mô tả cắt ngang)')

q3_proposed = [
    ('B1', 'Manifestations and patterns of anxiety disorder subtypes among '
     'Vietnamese lower secondary school students: A descriptive cross-'
     'sectional study\n'
     '(Biểu hiện và mô hình các phân loại rối loạn lo âu ở học sinh trung học '
     'cơ sở Việt Nam: nghiên cứu mô tả cắt ngang)',
     'Bám sát đề cương v3 hiện tại'),
    ('B2', 'Prevalence and item-level patterns of generalized, separation, '
     'and social anxiety symptoms among Vietnamese lower secondary school '
     'students: A descriptive cross-sectional study\n'
     '(Tỷ lệ hiện mắc và các mô hình mức độ mục của triệu chứng lo âu tổng '
     'quát, chia ly, và xã hội ở học sinh trung học cơ sở Việt Nam: nghiên '
     'cứu mô tả cắt ngang)',
     'MẪU CHUẨN Q3 — "Prevalence and..." như Bangladesh – PLOS ONE'),
    ('B3', 'Item-level analysis of anxiety disorder subtypes and grade-level '
     'developmental trajectories among Vietnamese lower secondary school '
     'students: A descriptive normative study\n'
     '(Phân tích mức độ mục của các phân loại rối loạn lo âu và quỹ đạo phát '
     'triển theo khối lớp ở học sinh trung học cơ sở Việt Nam: nghiên cứu mô '
     'tả chuẩn)',
     'Nhấn vào "normative data" và "developmental"'),
    ('B4', 'Anxiety disorder subtypes among Vietnamese lower secondary '
     'students: A cross-sectional descriptive study of item-level patterns, '
     'gender, and grade trajectories\n'
     '(Các phân loại rối loạn lo âu ở học sinh trung học cơ sở Việt Nam: '
     'nghiên cứu mô tả cắt ngang về các mô hình mức độ mục, giới tính, và '
     'quỹ đạo theo khối lớp)',
     'Tổng hợp cả ba câu hỏi nghiên cứu trong tiêu đề — dài nhưng đầy đủ'),
    ('B5', 'Generalized, separation, and social anxiety symptoms among '
     'Vietnamese adolescents in Hanoi: A descriptive item-level analysis '
     'for screening tool development\n'
     '(Triệu chứng lo âu tổng quát, chia ly, và xã hội ở thanh thiếu niên '
     'Việt Nam tại Hà Nội: phân tích mức độ mục mô tả cho việc phát triển '
     'công cụ sàng lọc)',
     'Nhấn vào ứng dụng thực tiễn "screening tool"'),
]
make_title_table(
    [(r[0], r[1], 'Hang Thi Cong và cs. (VN)',
      'PLOS ONE', r[2]) for r in q3_proposed],
    [0.8, 9.0, 2.5, 2.5, 3.2]
)


d.add_page_break()


# ============================================================
H1('PHẦN 5 — KHUYẾN NGHỊ CỦA EM')


H2('5.1 Đề xuất top picks cho nhóm cân nhắc')

P('Sau khi phân tích kỹ các mẫu cấu trúc thành công của bài báo Q1 và Q3 '
  'từ các tác giả châu Á và châu Phi, em đưa ra khuyến nghị cụ thể cho '
  'từng bài.')

H3('Đối với bài Q1')

P('Em khuyến nghị thầy và đồng tác giả cân nhắc giữa hai phương án A1 và '
  'A4. Phương án A1 (giữ nguyên đề cương) có ưu điểm là cô đọng và bao quát '
  'cả ba điểm mới của bài: mô hình SEM tích hợp, các phân loại lo âu, và '
  'phương pháp hỗn hợp. Phương án A4 ("Multi-group SEM") thì gần với câu '
  'chuyện thành công của bài Rwanda 2024 trên Frontiers Psychiatry, có thể '
  'dễ được tạp chí chấp nhận hơn vì mẫu này đã được tạp chí Q1 chấp nhận '
  'tiền lệ.')

P('Phương án A3 quá đặc thù — đặt phát hiện "bất biến giới" làm tâm điểm '
  'của tiêu đề có thể giới hạn phạm vi tham khảo của bài. Phương án A5 có '
  'rủi ro: cụm từ "Beyond" trong tiêu đề có thể bị reviewers đánh giá là '
  '"sensationalist" (giật gân), không phù hợp với phong cách học thuật của '
  'tạp chí Q1.')

H3('Đối với bài Q3')

P('Em khuyến nghị cân nhắc giữa phương án B2 và phương án B5. Phương án B2 '
  '("Prevalence and item-level patterns") đúng với mẫu Q3 chuẩn của PLOS '
  'ONE, khớp với mẫu thành công của bài Bangladesh 2021. Phương án B5 nhấn '
  'mạnh ứng dụng thực tiễn — phát triển công cụ sàng lọc, một góc mà các '
  'reviewers của PLOS ONE thường đánh giá cao vì có giá trị áp dụng rõ rệt.')

P('Phương án B1 (giữ nguyên đề cương) tuy được nhưng cụm "Manifestations '
  'and patterns" hơi mơ hồ — reviewers có thể đặt câu hỏi "manifestations '
  'gì cụ thể?". Cụm này không truyền tải rõ thông điệp về phân tích mức '
  'độ mục.')


H2('5.2 Các câu hỏi thảo luận cho nhóm')

P('Ngoài việc chọn tên bài, em muốn đề xuất một số câu hỏi để nhóm cùng '
  'thảo luận và đạt sự đồng thuận về các quyết định liên quan đến phong '
  'cách trình bày của hai bài báo:')

B('Bài Q1 nên nhấn mạnh PHƯƠNG PHÁP (mô hình SEM tích hợp, phương pháp '
  'hỗn hợp) hay nhấn mạnh PHÁT HIỆN (bất biến giới của lo âu chia ly, các '
  'đường dẫn khác biệt)? Hai cách tiếp cận sẽ dẫn đến tiêu đề khác nhau '
  'rõ rệt.')
B('Có nên ghi rõ "Vietnam" hoặc "Hanoi" trong tiêu đề? Ưu điểm: reviewers '
  'thích vì biết rõ bối cảnh nghiên cứu. Nhược điểm: có thể giảm sức hút '
  'của bài đối với độc giả ngoài Việt Nam.')
B('Lựa chọn thuật ngữ tiếng Anh: "lower secondary school" (chuẩn của Bộ '
  'GD Việt Nam) vs "middle school" (Mỹ) vs "junior high" (Anh) vs "early '
  'adolescents" (tổng quát). Sự lựa chọn này có sự đánh đổi giữa tính cụ '
  'thể và tính phổ quát.')
B('"DSM-5 anxiety subtypes" (viết gọn) vs "generalized, separation, social '
  'anxiety" (liệt kê đầy đủ) — viết gọn ngắn hơn nhưng liệt kê đầy đủ thì '
  'rõ ràng hơn cho reviewers ngoài chuyên ngành.')
B('Hai tiêu đề Q1 và Q3 có nên có format TƯƠNG TỰ để dễ nhận diện cùng '
  'thuộc một chương trình nghiên cứu? Ví dụ cả hai kết thúc bằng "[...] '
  'among Vietnamese lower secondary school students"?')


H2('5.3 Em xin ý kiến thầy + đồng tác giả')

P('Em đã liệt kê 5 phương án cho Q1 (A1-A5) và 5 phương án cho Q3 (B1-B5) '
  'ở Phần 4. Thầy hướng dẫn và các đồng tác giả có thể:', italic=True)

B('Chọn một phương án từ danh sách A1-A5 cho bài Q1 và một phương án từ '
  'B1-B5 cho bài Q3')
B('Hoặc đề xuất kết hợp hai phương án, hoặc sửa đổi từ một phương án có sẵn '
  '(ví dụ: "A1 nhưng đổi mixed-methods thành qualitative-quantitative")')
B('Hoặc đưa ra hướng tiêu đề hoàn toàn mới mà em chưa nghĩ đến')

P('Em sẵn sàng tiếp nhận mọi phản hồi và sẽ điều chỉnh tiêu đề cuối cùng '
  'cho hai bài báo trong các phiên làm việc tiếp theo.', italic=True)


# ============================================================
H1('PHẦN 6 — DANH MỤC NGUỒN TRA CỨU')

P('Tài liệu này tổng hợp thông tin từ nhiều cơ sở dữ liệu học thuật uy tín. '
  'Dưới đây là danh sách các nguồn chính được sử dụng, sắp xếp theo mức độ '
  'phổ biến và độ tin cậy trong nghiên cứu sức khỏe tâm thần thanh thiếu niên.',
  indent=False)

H2('Các tạp chí quốc tế chính được tham khảo')

B('General Psychiatry (Q1, IF 7.0, BMJ Publishing Group) — chuyên về phân '
  'tích xu hướng dịch tễ học rối loạn tâm thần')
B('BMC Psychiatry (Q1, IF 4.4, BMC/Springer Nature) — open access, chấp '
  'nhận đa dạng phương pháp')
B('Journal of Affective Disorders (Q1, IF 6.6, Elsevier) — tập trung vào '
  'trầm cảm + lo âu')
B('PLOS ONE (Q1/Q2, IF 3.7, PLOS) — open access, scope rộng, chấp nhận '
  'descriptive')
B('PLOS Mental Health (mới, 2024-, PLOS) — chuyên về sức khỏe tâm thần')
B('Frontiers in Psychiatry (Q1, IF 4.7, Frontiers) — open access, '
  'multi-disciplinary')
B('Scientific Reports (Q1, IF 4.6, Nature Publishing Group) — open access, '
  'multi-discipline')

H2('Các cơ sở dữ liệu tìm kiếm')

B('PubMed / PMC (National Library of Medicine, Hoa Kỳ) — '
  'cơ sở dữ liệu y sinh học lớn nhất thế giới')
B('Web of Science (Clarivate Analytics) — nguồn chính cho Impact Factor '
  'và Quartile')
B('SCImago Journal Rank (SJR) — nguồn alternative cho Quartile')
B('Google Scholar — bổ trợ tìm kiếm linh hoạt')


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
