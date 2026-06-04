# -*- coding: utf-8 -*-
"""Dịch A11 — Long An 2025 — Trần Đức Sĩ — TC Y Dược học Phạm Ngọc Thạch"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.59715/pntjmp.4.4.4', size=10)
add_heading(doc, 'Tỷ lệ trầm cảm, lo âu, stress ở học sinh trung học phổ thông tại thành phố Tân An, tỉnh Long An', 1)
h = doc.add_paragraph()
r = h.add_run('The prevalence of depression, anxiety, stress in high school students in Tan An city, Long An province')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề', 'Tỷ lệ trầm cảm, lo âu, stress ở học sinh THPT tại TP Tân An, tỉnh Long An'),
    ('Tác giả', 'Trần Đức Sĩ (1), Mai Phương Dung (1)'),
    ('Cơ quan', '(1) Bộ môn Y học Gia đình, Khoa Y, Trường ĐH Y khoa Phạm Ngọc Thạch'),
    ('Tạp chí', 'Tạp chí Y Dược học Phạm Ngọc Thạch, 2025, 4(4): 31–39'),
    ('DOI', '10.59715/pntjmp.4.4.4'),
    ('Loại NC', 'Cắt ngang mô tả'),
    ('Mẫu', '360 HS THPT tại 3 trường TP Tân An, Long An: THPT Chuyên Long An + THPT Lê Quý Đôn (công lập) + THPT Hà Long (tư thục); 1–3/2024'),
    ('Công cụ', 'DASS-21 (trầm cảm + lo âu + stress)'),
])

add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'TC Y Dược học Phạm Ngọc Thạch — có DOI, peer-reviewed VN. ĐH Y khoa PNT uy tín.',
    'n = 360 — 3 LOẠI TRƯỜNG: chuyên + công lập + tư thục → SO SÁNH hiếm.',
    'Lo âu 57,20% (DASS-21) — CAO. Phù hợp dải DASS-21 VN (40–86%).',
    'Trường CHUYÊN + TƯ có SKTT kém hơn trường THƯỜNG — phát hiện QUAN TRỌNG.',
    'Năm CUỐI CẤP (lớp 12) tỷ lệ CAO NHẤT — phù hợp VN21 (stress tăng lớp 6→9).',
    'Long An — Miền Tây Nam Bộ, cầu nối TPHCM–ĐBSCL. Ít NC SKTT.',
    'DASS-21 — so sánh được với Hoa 2024, Bảo Quyên 2025, An Giang 2025, Hải Phòng 2024.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'n = 360 — trung bình. Chỉ 3 trường TP Tân An.',
    'DASS-21 sàng lọc — không chẩn đoán.',
    'TC VN — không PubMed/Scopus.',
    'Không phân tích yếu tố liên quan (hồi quy) — chỉ mô tả tỷ lệ.',
    'Loại trừ HS có hoàn cảnh đặc biệt (tang, bất hòa, chia tay) — có thể bỏ sót nhóm nguy cơ cao.',
]:
    add_p(doc, f'• {b}')

add_heading(doc, 'KẾT QUẢ CHÍNH', 2)

add_heading(doc, 'Bảng 1. Tỷ lệ DASS-21 — Tổng và theo loại trường (n = 360)', 3)
add_table(doc,
    ['Chỉ số', 'Tổng', 'Trường chuyên', 'Công lập', 'Tư thục', 'So sánh'],
    [['Trầm cảm', '38,48%', 'Cao hơn', 'Thấp hơn', 'Cao hơn', 'Chuyên + Tư > Thường'],
     ['Lo âu', '57,20%', 'Cao hơn', 'Thấp hơn', 'Cao hơn', 'Chuyên + Tư > Thường'],
     ['Stress', '33,37%', 'Cao hơn', 'Thấp hơn', 'Cao hơn', 'Chuyên + Tư > Thường']],
    widths=[2.0, 2.0, 2.5, 2.5, 2.5, 3.0])
add_p(doc, 'Trường CHUYÊN: áp lực thi cử + kỳ vọng cao → SKTT kém hơn. Trường TƯ: có thể do áp lực tài chính gia đình + cạnh tranh. Trường THƯỜNG: ít áp lực → SKTT tốt hơn.', size=9, italic=True)

add_heading(doc, 'Bảng 2. Tỷ lệ theo năm học', 3)
add_table(doc,
    ['Khối lớp', 'Xu hướng', 'Ghi chú'],
    [['Lớp 10', 'Thấp nhất', ''],
     ['Lớp 11', 'Trung bình', ''],
     ['Lớp 12', 'CAO NHẤT', 'Áp lực thi ĐH + tốt nghiệp']],
    widths=[3.0, 3.0, 5.0])
add_p(doc, 'Phù hợp VN21 (Trần Thảo Vi 2024): stress ESSA tăng từ lớp 6 đến lớp 9 (β=4,73 cho học thêm). Lớp 12 = năm thi ĐH.', size=9, italic=True)

add_p(doc, 'Đối chiếu liên bài', bold=True)
add_p(doc, 'Lo âu 57,20% — so với: Hoa 2024 VN01 40,6% (GAD-7, khác công cụ); Bảo Quyên VN16 86,2% (DASS-21, Hà Nội); An Giang VN18 61,2% (DASS-21); Thảo Vi VN19 65,8% (DASS-21, Huế); Vĩnh Lộc VN20 25,1% (DASS-Y, TPHCM); Hải Phòng A10 53,81% (DASS-21). Long An nằm ở mức TRUNG BÌNH–CAO trong dải DASS-21 VN.')
add_p(doc, 'Trường chuyên cao hơn — phù hợp Norway QT21 (bất mãn trường giải thích chính), Wen QT08 (áp lực OR=11,58). Áp lực trường chuyên ≈ áp lực thi chuyển cấp.')
add_p(doc, 'Năm cuối cấp cao nhất — phù hợp VN21 (stress tăng lớp 6→9, thi vào lớp 10). Tại đây là thi tốt nghiệp + ĐH.')

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'TC Y Dược học PNT có DOI. ĐH Y khoa PNT — uy tín TPHCM.',
    'So sánh 3 LOẠI TRƯỜNG (chuyên/công lập/tư thục) — HIẾM trong NC VN. Phát hiện trường chuyên + tư > thường.',
    'DASS-21 — so sánh được. Long An Miền Tây — mở rộng vùng.',
    'Tỷ lệ theo khối lớp — xác nhận năm cuối cao nhất.',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'n = 360 — nhỏ so với NC khác (Hoa VN01 3.910; A8 Medicine 5.730).',
    'KHÔNG phân tích yếu tố liên quan (hồi quy) — chỉ mô tả. Không biết yếu tố nào DỰ BÁO lo âu.',
    'Loại trừ HS hoàn cảnh đặc biệt — bỏ sót nhóm nguy cơ CAO NHẤT.',
    'TC VN — không PubMed/Scopus.',
    'Chỉ TP Tân An — không đại diện toàn Long An hoặc ĐBSCL.',
    'Không đo: MXH, giấc ngủ, bắt nạt, áp lực HT cụ thể (ESSA) — yếu tố quan trọng.',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'Cần NC yếu tố liên quan (hồi quy đa biến) tại Long An/ĐBSCL — hiện chỉ mô tả tỷ lệ.',
    'So sánh chi tiết: trường chuyên vs công lập vs tư — TÁCH yếu tố áp lực cụ thể.',
    'Thêm ESSA (áp lực HT) + GAD-7 (lo âu riêng) — đo đa công cụ trên cùng mẫu.',
    'Mở rộng: so sánh Tân An (đô thị) vs huyện nông thôn Long An — khác biệt vùng.',
    'NC can thiệp: giảm áp lực tại trường chuyên — BMC NMA QT29 (CBT SUCRA 0,66), B8 Sri Lanka (cluster RCT GV cung cấp).',
]:
    add_red(doc, f'• {s}')

add_p(doc, 'Đánh giá: ⭐⭐⭐ Trung bình. TC VN DOI, n=360, so sánh 3 loại trường. Nhưng chỉ mô tả, thiếu hồi quy, n nhỏ.', bold=True)

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '47_LongAn_2025.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['360', '57,20', '38,48', '33,37', 'DASS-21', 'Tân An', 'Long An', 'chuyên']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Numbers: {ok}/{len(checks)}')
