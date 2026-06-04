"""Build CHUA_gia_thuyet_NC_3_thanh_8_chuong_1.docx
Chua 3 gia thuyet cua tro ly CTH thanh 8 gia thuyet chinh xac hon, co co so y van.
Style: CTH v6.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/CHUA_gia_thuyet_NC_chuong_1.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = BLACK

def para(text, indent=True, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(13)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def add_table(header, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

def ref_entry(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CHỮA HỆ THỐNG GIẢ THUYẾT NGHIÊN CỨU\n(Chương 1 — Tổng quan tài liệu và giả thuyết)')
r.bold = True; r.font.size = Pt(14)
para('')

# ============================================================
# 1. Phân tích 3 giả thuyết hiện tại
# ============================================================
H('1. Đánh giá ba giả thuyết hiện tại', level=2)

para(
    'Ba giả thuyết do trợ lý nghiên cứu đề xuất ban đầu phản ánh ĐÚNG bốn mảng '
    'kết quả chính của chương 3 — mức độ rối loạn lo âu, khác biệt nhân khẩu, '
    'cường độ yếu tố nguy cơ–bảo vệ. Tuy nhiên, ba giả thuyết này còn ba điểm '
    'cần chỉnh sửa để đáp ứng chuẩn falsifiability của Popper (1959) và chuẩn '
    'a priori prediction của thiết kế nghiên cứu định lượng theo Creswell '
    '(2014).'
)

para(
    'Thứ nhất, các giả thuyết quá RỘNG — chỉ nêu "có sự khác biệt" mà không '
    'chỉ rõ HƯỚNG khác biệt. Nói cách khác, giả thuyết "có khác biệt giới '
    'tính" sẽ luôn được xác nhận miễn là p < 0,05 ở bất kỳ chiều nào, không '
    'cho phép kiểm chứng thực sự một dự đoán cụ thể. Thứ hai, ba giả thuyết '
    'thiếu cơ sở y văn DẪN DẮT — không nêu nguồn nào cho phép phát biểu giả '
    'thuyết theo hướng đó. Thứ ba, ba giả thuyết BỎ SÓT các phát hiện QUAN '
    'TRỌNG khác trong chương 3 — biện pháp đối phó (BPĐP), khoảng trống giao '
    'tiếp cha mẹ–con, và hiện tượng tự trọng có cường độ ngang bằng áp lực '
    'học tập.'
)

# Bảng đánh giá 3 giả thuyết
H('1.1. Bảng đánh giá chi tiết ba giả thuyết hiện tại', level=3)
caption('Bảng 1. Đánh giá ba giả thuyết do trợ lý nghiên cứu đề xuất ban đầu')
add_table(
    ['Giả thuyết', 'Đánh giá', 'Vấn đề chính'],
    [
        ['H1: Có sự khác biệt về mức độ giữa ba dạng RLLA',
         'Đúng hướng — kết quả xác nhận',
         'Quá rộng. Không nêu HƯỚNG: dạng nào cao nhất, dạng nào thấp nhất.'],
        ['H2: Có khác biệt có ý nghĩa thống kê về RLLA theo giới và khối lớp',
         'Đúng hướng — kết quả xác nhận một phần',
         'Không nêu HƯỚNG (nữ > nam? lớp lớn > lớp nhỏ?). KHÔNG bao gồm trường hợp ngoại lệ — lo âu chia ly KHÔNG khác biệt theo giới.'],
        ['H3: Yếu tố nguy cơ MẠNH HƠN yếu tố bảo vệ',
         'BẤT THƯỜNG về mặt phương pháp luận',
         'So sánh "tổng cường độ" hai nhóm yếu tố là khó kiểm định trước. Bỏ sót CHIỀU tác động của từng yếu tố cụ thể (ALHT dương; TTr âm; HTCM âm...). Phù hợp post-hoc, không phù hợp a priori.'],
    ]
)
para('')

# ============================================================
# 2. Đề xuất 8 giả thuyết chỉnh sửa
# ============================================================
H('2. Đề xuất tám giả thuyết chỉnh sửa', level=2)

para(
    'Trên cơ sở kết quả nghiên cứu thực tế của chương 3 và y văn hỗ trợ, đề '
    'xuất tám giả thuyết được tổ chức theo bốn nhóm — mức độ và biểu hiện, '
    'khác biệt nhân khẩu, yếu tố ảnh hưởng, và biện pháp đối phó. Mỗi giả '
    'thuyết nêu rõ HƯỚNG dự đoán và CƠ SỞ y văn dẫn dắt — đáp ứng chuẩn a '
    'priori prediction theo Creswell (2014).'
)

H('2.1. Nhóm giả thuyết về mức độ và biểu hiện rối loạn lo âu', level=3)

para(
    'Giả thuyết H1 — Trong ba dạng rối loạn lo âu được đo lường ở học sinh '
    'trung học cơ sở Việt Nam, lo âu LAN TỎA có cường độ cao nhất, tiếp '
    'theo là lo âu XÃ HỘI, và thấp nhất là lo âu CHIA LY. Cơ sở: tổng quan '
    'của Beesdo, Knappe và Pine (2009) về phát triển lo âu ở vị thành niên '
    '— Generalized Anxiety và Social Anxiety tăng từ tuổi 11 trong khi '
    'Separation Anxiety đã qua đỉnh ở tuổi 7–9. Phù hợp với khung '
    'developmental psychopathology — lứa tuổi 11–15 đang chuyển dịch giai '
    'đoạn lo âu chính.'
)

para(
    'Giả thuyết H2 — Mệnh đề liên quan đến THẤT BẠI HỌC TẬP và áp lực '
    'thành tích có điểm trung bình CAO HƠN các mệnh đề khác trong thang '
    'đo lo âu lan tỏa. Cơ sở: Khảo sát PISA 2015 của OECD ghi nhận 66% '
    'học sinh lo về điểm kém và 55% lo về kiểm tra dù chuẩn bị tốt '
    '(Pascoe, Hetrick & Parker, 2020); văn hóa giáo dục Á Châu nhấn mạnh '
    'perfectionism (Wen và cộng sự, 2020).'
)

H('2.2. Nhóm giả thuyết về khác biệt nhân khẩu học', level=3)

para(
    'Giả thuyết H3a — Học sinh nữ có điểm rối loạn lo âu lan tỏa và lo âu '
    'xã hội CAO HƠN học sinh nam. Cơ sở: tổng quan hệ thống của McLean, '
    'Asnaani, Litz và Hofmann (2011) cùng phân tích tổng hợp của Salk, '
    'Hyde và Abramson (2017) khẳng định nữ > nam ở GAD và Social Anxiety '
    'với tỷ số nguy cơ ~1,5–2 lần, chênh lệch mở rộng sau dậy thì.'
)

para(
    'Giả thuyết H3b — Khác biệt giới tính KHÔNG xuất hiện ở lo âu chia ly. '
    'Cơ sở: Allen, Lavallee, Herren, Ruhe và Schneider (2010) cho thấy '
    'Separation Anxiety Disorder ít chịu ảnh hưởng giới hơn so với GAD '
    'và Social Anxiety ở vị thành niên — cơ chế sinh học liên quan hệ '
    'thống gắn bó (attachment system), vốn ít biến đổi theo giới.'
)

para(
    'Giả thuyết H4 — Lo âu lan tỏa và lo âu xã hội TĂNG dần theo khối lớp '
    '(lớp 6 → lớp 9), trong khi lo âu chia ly GIẢM dần. Cơ sở: '
    'developmental psychopathology của Beesdo, Knappe và Pine (2009) cùng '
    'mô hình của Rapee và Spence (2004) về phát triển lo âu xã hội theo '
    'dậy thì khi tự nhận thức xã hội tăng. Nói cách khác, ba dạng lo âu '
    'có quỹ đạo phát triển khác nhau theo tuổi — phù hợp với khung '
    'sinh-tâm-xã hội.'
)

H('2.3. Nhóm giả thuyết về yếu tố ảnh hưởng', level=3)

para(
    'Giả thuyết H5 — Trong các yếu tố nguy cơ, áp lực học tập có cường '
    'độ TÁC ĐỘNG MẠNH NHẤT đến rối loạn lo âu, với hệ số chuẩn hóa β > '
    '0,4. Cơ sở: Wen và cộng sự (2020) tại Trung Quốc nông thôn ghi '
    'nhận OR = 11,58 (KTC 95% từ 4,16 đến 32,19) cho áp lực rất cao; '
    'Trần Thảo Vi và cộng sự (2024) tại Huế ghi nhận β = 4,73 cho học '
    'thêm; tổng quan của Pascoe, Hetrick và Parker (2020) khẳng định áp '
    'lực học tập là yếu tố nguy cơ chính ở vị thành niên Á Châu.'
)

para(
    'Giả thuyết H6 — Tự trọng có cường độ tác động bảo vệ NGANG BẰNG với '
    'cường độ tác động nguy cơ của áp lực học tập, không yếu hơn như giả '
    'định thông thường. Cơ sở: phân tích tổng hợp của Cai và cộng sự '
    '(2025) trên các thử nghiệm ngẫu nhiên có đối chứng can thiệp '
    'resilience tại trường — xác định tự trọng là thành phần trung tâm '
    'của resilience cùng với kỹ năng ứng phó, kết nối xã hội và lạc quan; '
    'khung resilience toàn cầu của Masten (2014) khẳng định tự trọng là '
    'một trong các "ordinary magic" — yếu tố nội tại bảo vệ trẻ em trước '
    'nghịch cảnh. Phát hiện này có hàm ý CRITICAL — chiến lược can thiệp '
    'không nên tập trung đơn trục vào giảm áp lực mà cần SONG TRỤC kết '
    'hợp tăng tự trọng.'
)

para(
    'Giả thuyết H7 — Khoảng trống giao tiếp giữa cha mẹ và con là rào '
    'cản chính trong việc phát huy hiệu ứng bảo vệ của hỗ trợ gia đình; '
    'cụ thể, mức độ "hỗ trợ cha mẹ cảm nhận" cao nhưng "khả năng chia '
    'sẻ với gia đình" THẤP HƠN rõ rệt. Cơ sở: Khảo sát Sức khỏe Tâm '
    'thần Vị thành niên Việt Nam (V-NAMHS) ghi nhận chỉ 5,1% phụ huynh '
    'xác định được con cần trợ giúp tâm lý (UNICEF Việt Nam, 2022) — '
    'phản ánh khoảng cách hệ thống giữa ý định hỗ trợ và khả năng nhận '
    'biết nhu cầu của con.'
)

H('2.4. Nhóm giả thuyết về biện pháp đối phó', level=3)

para(
    'Giả thuyết H8 — TẦN SUẤT sử dụng biện pháp đối phó KHÔNG đồng '
    'nghĩa với CHẤT LƯỢNG đối phó; do đó, tăng tần suất KHÔNG tự động '
    'làm GIẢM rối loạn lo âu, và ở một số trường hợp có thể đi đôi với '
    'lo âu cao hơn. Cơ sở: tổng quan của Compas, Jaser, Bettis và cộng '
    'sự (2017) trong Psychological Bulletin về hiện tượng MALADAPTIVE '
    'COPING ESCALATION — học sinh càng lo âu càng dùng nhiều chiến '
    'lược đối phó, nhưng nếu chiến lược không hiệu quả thì lo âu vẫn '
    'duy trì. Phù hợp với mô hình 14 nhân tố của Brief-COPE (Carver, '
    '1997) phân biệt rõ chiến lược adaptive (problem-focused, '
    'support-seeking) với maladaptive (avoidance, self-blame, '
    'rumination).'
)

# ============================================================
# 3. So sánh trực tiếp
# ============================================================
H('3. So sánh trực tiếp giả thuyết cũ và giả thuyết chỉnh sửa', level=2)

caption('Bảng 2. Đối chiếu ba giả thuyết hiện tại với tám giả thuyết chỉnh sửa')
add_table(
    ['Giả thuyết hiện tại', 'Giả thuyết chỉnh sửa', 'Cải thiện chính'],
    [
        ['H1 (cũ): Có khác biệt giữa 3 dạng RLLA',
         'H1 (mới): Lan tỏa > xã hội > chia ly',
         'Nêu HƯỚNG cụ thể, dẫn Beesdo 2009'],
        ['—',
         'H2 (mới): Mệnh đề thất bại học tập có ĐTB cao nhất',
         'Bổ sung — chi tiết phát hiện cụ thể'],
        ['H2 (cũ): Có khác biệt theo giới và khối',
         'H3a + H3b + H4 (mới)',
         'Tách thành ba giả thuyết riêng, nêu HƯỚNG, ghi nhận ngoại lệ'],
        ['H3 (cũ): Nguy cơ mạnh hơn bảo vệ',
         'H5 + H6 + H7 (mới)',
         'Thay so sánh tổng cường độ bằng nêu yếu tố cụ thể có |β| > 0,4'],
        ['—',
         'H8 (mới): Biện pháp đối phó tăng KHÔNG = lo âu giảm',
         'Bổ sung — phản ánh phát hiện chương 3.5 + Compas 2017'],
    ]
)
para('')

para(
    'Bảng 2 cho thấy ba giả thuyết hiện tại được mở rộng thành tám giả '
    'thuyết — tăng tính chính xác (nêu hướng), bao quát đầy đủ phát hiện '
    '(thêm BPĐP và khoảng trống giao tiếp), và đáp ứng chuẩn falsifiability '
    '(mỗi giả thuyết có thể bị bác bỏ bằng dữ liệu cụ thể). Phù hợp với '
    'nguyên tắc thiết kế nghiên cứu định lượng theo Creswell (2014) — giả '
    'thuyết cần được phát biểu RÕ RÀNG, có thể KIỂM ĐỊNH, và DẪN dắt từ '
    'tổng quan tài liệu trước đó.'
)

# ============================================================
# 4. Khung lý thuyết tổng hợp
# ============================================================
H('4. Khung lý thuyết tích hợp tám giả thuyết', level=2)

para(
    'Tám giả thuyết có thể được tổ chức theo khung biopsychosocial-'
    'developmental của Beesdo, Knappe và Pine (2009) kết hợp khung '
    'risk-resilience của Masten (2014). Khung này cho phép kiểm tra '
    'đồng thời ba lớp tác động — đặc điểm nhân khẩu (giới, tuổi), '
    'yếu tố môi trường (gia đình, trường học, bạn bè), và đặc điểm cá '
    'nhân (tự trọng, biện pháp đối phó) — trong một mô hình SEM duy '
    'nhất.'
)

para(
    'Cấu trúc này thuyết phục hơn cấu trúc mô tả thuần túy ("nghiên cứu '
    'thực trạng và yếu tố ảnh hưởng") vì cho phép người đọc theo dõi '
    'mạch lý thuyết xuyên suốt — chương 1 đặt câu hỏi và giả thuyết, '
    'chương 2 mô tả phương pháp đo lường để kiểm chứng, chương 3 báo '
    'cáo kết quả kiểm chứng từng giả thuyết, và chương 4 thảo luận '
    'hàm ý cùng đề xuất can thiệp. Nói cách khác, cấu trúc giả thuyết '
    'rõ ràng là CƠ SỞ logic đảm bảo tính chặt chẽ của toàn luận án.'
)

# ============================================================
# 5. Đề xuất paste vào chương 1
# ============================================================
H('5. Đề xuất hình thức trình bày trong chương 1', level=2)

para(
    'Tám giả thuyết trên có thể được paste vào mục "Giả thuyết nghiên '
    'cứu" của chương 1 — sau phần tổng quan tài liệu và trước phần mục '
    'tiêu nghiên cứu. Đề xuất tổ chức theo bốn tiểu mục:'
)
para('1.X.1. Giả thuyết về mức độ và biểu hiện (H1 + H2)', indent=False, justify=False)
para('1.X.2. Giả thuyết về khác biệt nhân khẩu (H3a + H3b + H4)', indent=False, justify=False)
para('1.X.3. Giả thuyết về yếu tố ảnh hưởng (H5 + H6 + H7)', indent=False, justify=False)
para('1.X.4. Giả thuyết về biện pháp đối phó (H8)', indent=False, justify=False)

para(
    'Trong mỗi tiểu mục, viết theo công thức ba câu — câu thứ nhất '
    'phát biểu giả thuyết, câu thứ hai nêu cơ sở y văn dẫn dắt, câu '
    'thứ ba nêu cách kiểm chứng (ví dụ "Giả thuyết được kiểm chứng '
    'bằng kiểm định t độc lập với mức ý nghĩa α = 0,05" hoặc "...bằng '
    'mô hình phương trình cấu trúc SEM với ngưỡng |β| ≥ 0,1").'
)

# ============================================================
# 6. TLTK
# ============================================================
H('6. Tài liệu tham khảo', level=2)

para('Tiếng Việt', indent=False, justify=False)
ref_entry('UNICEF Việt Nam, Bộ Lao động – Thương binh và Xã hội, và Tổng cục Thống kê. (2022). Khảo sát Sức khỏe Tâm thần Vị thành niên Việt Nam (V-NAMHS 2022). Hà Nội.')
ref_entry('Trần, T. V., và cộng sự. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine.')

para('Tiếng Anh', indent=False, justify=False)
ref_entry('Allen, J. L., Lavallee, K. L., Herren, C., Ruhe, K., & Schneider, S. (2010). DSM-IV criteria for childhood separation anxiety disorder: Informant, age, and sex differences. Journal of Anxiety Disorders, 24(8), 946–952. https://doi.org/10.1016/j.janxdis.2010.06.022')
ref_entry('Beesdo, K., Knappe, S., & Pine, D. S. (2009). Anxiety and anxiety disorders in children and adolescents: Developmental issues and implications for DSM-V. Psychiatric Clinics of North America, 32(3), 483–524. https://doi.org/10.1016/j.psc.2009.06.002')
ref_entry('Cai, S., et al. (2025). Resilience as a protective factor against anxiety in adolescents. Frontiers in Psychiatry.')
ref_entry('Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100. https://doi.org/10.1207/s15327558ijbm0401_6')
ref_entry('Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991.')
ref_entry('Creswell, J. W. (2014). Research design: Qualitative, quantitative, and mixed methods approaches (4th ed.). SAGE Publications.')
ref_entry('Masten, A. S. (2014). Global perspectives on resilience in children and youth. Child Development, 85(1), 6–20. https://doi.org/10.1111/cdev.12205')
ref_entry('McLean, C. P., Asnaani, A., Litz, B. T., & Hofmann, S. G. (2011). Gender differences in anxiety disorders: Prevalence, course of illness, comorbidity and burden of illness. Journal of Psychiatric Research, 45(8), 1027–1035.')
ref_entry('Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112.')
ref_entry('Popper, K. R. (1959). The logic of scientific discovery. Hutchinson.')
ref_entry('Rapee, R. M., & Spence, S. H. (2004). The etiology of social phobia: Empirical evidence and an initial model. Clinical Psychology Review, 24(7), 737–767.')
ref_entry('Salk, R. H., Hyde, J. S., & Abramson, L. Y. (2017). Gender differences in depression in representative national samples: Meta-analyses of diagnoses and symptoms. Psychological Bulletin, 143(8), 783–822.')
ref_entry('Wen, X., et al. (2020). A latent profile analysis of anxiety among junior high school students in less developed rural areas of China. International Journal of Environmental Research and Public Health, 17(11), 4079.')

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
