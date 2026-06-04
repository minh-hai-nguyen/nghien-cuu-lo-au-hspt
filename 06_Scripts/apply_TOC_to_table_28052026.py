# -*- coding: utf-8 -*-
"""Cap nhat bang TOC table dua tren heading JSON da extract.
28/05/2026."""
import os, sys, io, json, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
HEADINGS_FILE = os.path.join(ROOT, '06_Scripts', '_headings_28052026.json')

with open(HEADINGS_FILE, 'r', encoding='utf-8') as f:
    headings = json.load(f)
print(f"Loaded {len(headings)} headings from JSON")

d = Document(FILE)

# Find the TOC table (Table 0 — first cell 'MỞ ĐẦU')
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb
        break

if toc_table is None:
    print("KHONG TIM THAY bang TOC!")
    sys.exit(1)

print(f"TOC table: {len(toc_table.rows)} rows")

# Show existing TOC table content
print()
print("=== TOC table HIEN TAI ===")
existing_rows = []
for ri, row in enumerate(toc_table.rows):
    name = row.cells[0].text.strip()
    page = row.cells[1].text.strip() if len(row.cells) > 1 else ''
    existing_rows.append((ri, name, page))
    if ri < 10:
        print(f"  Row {ri}: {name!r} | trang {page}")

# Now compare with extracted headings
# Match each existing TOC row to a heading by text similarity
print()
print("=== SO SANH va MATCH ===")

def normalize(s):
    return re.sub(r'\s+', ' ', s.strip().lower())

# Build heading lookup: normalized text -> (level, page, idx)
heading_lookup = {}
for h in headings:
    key = normalize(h['text'])
    heading_lookup[key] = (h['level'], h['page'], h['idx'])
    # Also try without leading numbering
    no_num = re.sub(r'^[\d\.]+\s*', '', h['text']).strip()
    if no_num:
        heading_lookup[normalize(no_num)] = (h['level'], h['page'], h['idx'])

updated = 0
not_found = []
new_rows_data = []  # (name, new_page)

for ri, name, old_page in existing_rows:
    norm_name = normalize(name)
    no_num = re.sub(r'^[\d\.]+\s*', '', name).strip()
    norm_no_num = normalize(no_num)
    # Try exact match first, then no-num
    match = heading_lookup.get(norm_name) or heading_lookup.get(norm_no_num)
    if match:
        lv, pg, idx = match
        new_rows_data.append((ri, name, old_page, str(pg)))
        if str(pg) != old_page:
            updated += 1
    else:
        # Fuzzy: find heading containing this text or vice versa
        best = None
        for k, v in heading_lookup.items():
            if norm_no_num and (norm_no_num in k or k in norm_no_num):
                best = v
                break
        if best:
            lv, pg, idx = best
            new_rows_data.append((ri, name, old_page, str(pg)))
            if str(pg) != old_page:
                updated += 1
        else:
            new_rows_data.append((ri, name, old_page, '?'))
            not_found.append((ri, name))

print(f"Updated page numbers: {updated}")
print(f"Not found in document: {len(not_found)}")
if not_found:
    print(f"\nMUC KHONG TIM THAY trong document:")
    for ri, name in not_found[:10]:
        print(f"  Row {ri}: {name!r}")

# Show changes
print()
print("=== SO SANH old -> new (chi nhung dong THAY DOI) ===")
diff_count = 0
for ri, name, old_pg, new_pg in new_rows_data:
    if new_pg != '?' and str(old_pg) != new_pg:
        print(f"  Row {ri}: {name[:50]!r} | trang {old_pg} -> {new_pg}")
        diff_count += 1
        if diff_count > 20:
            print(f"  ... va {updated - 20} dong khac")
            break

# Apply: update cell[1] of each row with new page number
print()
print("=== APPLY UPDATES ===")
for ri, name, old_pg, new_pg in new_rows_data:
    if new_pg != '?' and str(old_pg) != new_pg:
        cell = toc_table.rows[ri].cells[1]
        # Clear all runs and add new page number
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''
        cell.paragraphs[0].text = ''
        new_run = cell.paragraphs[0].add_run(new_pg)
        new_run.font.name = 'Times New Roman'
        # Match existing size if was set
        new_run.font.size = Pt(13)

import shutil
backup = FILE.replace('.docx', '_BEFORE_TOC_PAGES_UPDATE.docx')
shutil.copy(FILE, backup)
print(f"Backup: {backup}")

d.save(FILE)
print(f"Saved: {FILE}")
print(f"Updated {updated} page numbers")
print(f"{len(not_found)} mục not matched (page giu nguyen)")
