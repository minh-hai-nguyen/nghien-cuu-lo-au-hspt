# -*- coding: utf-8 -*-
"""Bổ sung bảng số liệu vào 8 tóm tắt QT"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def rb(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(doc, t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); cell._tc.get_or_add_tcPr().append(s)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def rate(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

# ============================================================
# Rebuild all 8 QT summaries with tables
# ============================================================

# QT21 Norway
d = make_doc()
bl(d, 'Tóm tắt: Norway mental distress trends 2011\u20132024', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Giải thích khả thi cho xu hướng tăng căng thẳng tâm thần ở VTN Na Uy 2011\u20132024. Social Science & Medicine, 2025 (Q1, IF\u22485,4). 10 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Chỉ số', 'Phát hiện'],
    [['Xu hướng', 'Căng thẳng tâm thần tăng ở cả nam và nữ VTN Na Uy (2011\u20132024)'],
     ['Yếu tố giải thích mạnh nhất', 'Bất mãn với trường học giải thích phần lớn xu hướng tăng'],
     ['Yếu tố phụ', 'Thời gian dùng mạng xã hội giải thích một phần'],
     ['Phương pháp', 'Phân tích decomposition — phân tách đóng góp từng yếu tố'],
     ['Thiết kế', 'Phân tích xu hướng 13 năm, dữ liệu lặp lại']],
    widths=[4.0, 8.0])
d.add_paragraph()
rh2(d, 'Phản biện')
bl(d, 'Q1 IF=5,4, xu hướng 13 năm. Hạn chế: chỉ Na Uy, đo mental distress chung (không tách lo âu riêng).')
rh2(d, 'Hướng NC tiếp theo')
bl(d, 'So sánh xu hướng Na Uy với VN/châu Á. Can thiệp giảm bất mãn trường học.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50 Q1, xu hướng 13 năm, decomposition nâng cao.')
d.save('QT21_Norway_2025.docx'); sys.stderr.write('QT21 OK\n')

# QT22 Screen time
d = make_doc()
bl(d, 'Tóm tắt: Screen time & adolescent depression/anxiety (BJCP 2025)', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Mối liên quan cắt ngang và dọc giữa thời gian màn hình với trầm cảm và lo âu ở VTN. British J Clinical Psychology, 2025, 64:873\u2013887 (Q1). 15 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Thiết kế', 'Phát hiện chính', 'So sánh liên bài'],
    [['Cắt ngang', 'Screen time liên quan trầm cảm + lo âu', 'Phù hợp Chen 2023 (game OR=5,00)'],
     ['DỌC', 'Screen time DỰ BÁO trầm cảm/lo âu sau 1 năm', 'Vượt trội cắt ngang'],
     ['Ý nghĩa', 'Có thể suy luận chiều nhân quả', 'Hoàng Trung Học VN: Beta=0,176']],
    widths=[3.0, 5.0, 4.5])
d.add_paragraph()
rh2(d, 'Phản biện')
bl(d, 'Thiết kế DỌC — vượt trội. BJCP Q1. Hạn chế: cần đọc cỡ mẫu chi tiết từ PDF.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50 Thiết kế dọc, Q1, bằng chứng nhân quả.')
d.save('QT22_ScreenTime_2025.docx'); sys.stderr.write('QT22 OK\n')

# QT23 JAACAP
d = make_doc()
bl(d, 'Tóm tắt: US trends mental disorders 2013\u20132021 (JAACAP)', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Xu hướng rối loạn tâm thần ở trẻ em/VTN nhận điều trị trong hệ thống SKTT công lập Hoa Kỳ. JAACAP, 2024 (Q1, IF\u224811,0). N = 13.684.154 hồ sơ. 15 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Rối loạn', '2013', '2021', 'AOR (KTC 95%)', 'Xu hướng'],
    [['Lo âu', '9,6%', '19,2%', '2,17 (1,85\u20132,55)', '\u2191 Tăng gấp đôi'],
     ['Trầm cảm', '13,4%', '17,0%', '1,20 (1,03\u20131,41)', '\u2191 Tăng'],
     ['Sang chấn/stress', '22,7%', '27,4%', '1,31 (1,09\u20131,57)', '\u2191 Tăng'],
     ['Lưỡng cực', '\u2014', '\u2014', '\u2014', '\u2193 Giảm'],
     ['Rối loạn hành vi', '\u2014', '\u2014', '\u2014', '\u2193 Giảm']],
    widths=[3.0, 1.5, 1.5, 3.5, 2.5])
d.add_paragraph()
bl(d, 'Phát hiện nổi bật: Lo âu TĂNG GẤP ĐÔI (9,6% \u2192 19,2%) trong 8 năm — xu hướng mạnh nhất. So sánh: NSCH 2020 báo cáo lo âu tăng 61% (2016\u20132023) tại Mỹ.', bold=True)
rh2(d, 'Phản biện')
bl(d, 'JAACAP Q1 IF=11, n=13,7 triệu hồ sơ — dữ liệu quốc gia Mỹ. Hạn chế: chỉ trẻ ĐANG điều trị (thiên lệch chọn).')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 JAACAP Q1 IF=11, n kỷ lục 13,7 triệu, xu hướng 8 năm.')
d.save('QT23_JAACAP_US_2024.docx'); sys.stderr.write('QT23 OK\n')

# QT24 WHO Europe
d = make_doc()
bl(d, 'Tóm tắt: WHO Europe — SKTT trẻ em/thanh niên (Lancet 2025)', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Sức khỏe tâm thần của trẻ em và thanh niên trong khu vực WHO châu Âu. Lancet Regional Health Europe, 2025 (Q1, IF\u224815,0). 13 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Chỉ số', 'Dữ liệu'],
    [['VTN có rối loạn SKTT', '9 triệu (10\u201319 tuổi) tại châu Âu'],
     ['Lo âu + trầm cảm', 'Chiếm >50% tất cả ca rối loạn'],
     ['Xu hướng', 'Tỷ lệ tăng, đặc biệt sau COVID-19'],
     ['Giới tính', 'Nữ bị ảnh hưởng nhiều hơn'],
     ['Yếu tố nguy cơ', 'Mạng xã hội, áp lực học tập, nghèo đói'],
     ['Khuyến nghị', 'Tích hợp SKTT vào hệ thống giáo dục + y tế']],
    widths=[4.0, 8.0])
d.add_paragraph()
rh2(d, 'Phản biện')
bl(d, 'Lancet Q1 IF=15, phạm vi toàn châu Âu. Hạn chế: tổng quan chính sách, không NC gốc.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Lancet Q1 IF=15, tổng quan toàn diện cho chính sách.')
d.save('QT24_WHO_Europe_2025.docx'); sys.stderr.write('QT24 OK\n')

# QT26 UK NHS
d = make_doc()
bl(d, 'Tóm tắt: UK NHS Mental Health Statistics 2025', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Thống kê SKTT Anh: tỷ lệ, dịch vụ, tài trợ. House of Commons Library, UK Parliament, 2025. 46 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Chỉ số', '2014', '2025', 'Thay đổi'],
    [['Rối loạn TT phổ biến (16\u201324 tuổi)', '18,9%', '25,8%', '+36%'],
     ['Nữ 16\u201324', '\u2014', '36,1%', '\u2014'],
     ['Nam 16\u201324', '\u2014', '16,3%', 'Nữ gấp 2,2 lần'],
     ['Tự hại (suốt đời)', '\u2014', '10,3%', 'Tăng gấp 4 từ 2000'],
     ['Nữ 16\u201324 tự hại', '\u2014', '31,7%', 'Rất cao']],
    widths=[5.0, 2.0, 2.0, 3.0])
d.add_paragraph()
rh2(d, 'Phản biện')
bl(d, 'Nguồn chính thức Quốc hội Anh, NHS quốc gia. Hạn chế: báo cáo chính sách, không NC gốc. Chỉ Anh.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50 Nguồn chính thức, dữ liệu quốc gia, nhưng không peer-reviewed.')
d.save('QT26_UK_NHS_2025.docx'); sys.stderr.write('QT26 OK\n')

# QT27 Nature Social Media
d = make_doc()
bl(d, 'Tóm tắt: Social media in adolescents (Nature Human Behaviour 2025)', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Sử dụng mạng xã hội ở VTN có và không có rối loạn SKTT. Nature Human Behaviour, 9, 2025, pp. 1283\u20131299 (Q1, IF\u224824,0). N = 3.340, 11\u201319 tuổi, UK. 21 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Nhóm VTN', 'Hành vi mạng xã hội', 'Mức hài lòng'],
    [['Có rối loạn SKTT', 'Dùng MXH nhiều hơn đáng kể', 'Ít hài lòng hơn'],
     ['Không rối loạn SKTT', 'Dùng ít hơn', 'Hài lòng hơn'],
     ['Chẩn đoán', 'Bởi chuyên gia lâm sàng (không chỉ sàng lọc)', '\u2014'],
     ['48% rối loạn SKTT', 'Triệu chứng trước 18 tuổi', '\u2014']],
    widths=[4.0, 5.0, 3.0])
d.add_paragraph()
bl(d, 'So sánh: Chen 2023 (game OR=5,00), Hoàng Trung Học 2025 VN (điện tử Beta=0,176), Norway 2025 (MXH giải thích xu hướng tăng).')
rh2(d, 'Phản biện')
bl(d, 'Nature Q1 IF=24 — tạp chí TOP toàn cầu. n=3.340 quốc gia UK. Chẩn đoán lâm sàng. Hạn chế: cắt ngang, chỉ UK.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Nature Q1 IF=24, chẩn đoán lâm sàng, mẫu quốc gia.')
d.save('QT27_Nature_SocialMedia_2025.docx'); sys.stderr.write('QT27 OK\n')

# QT28 AJP Treatment
d = make_doc()
bl(d, 'Tóm tắt: Pediatric anxiety treatment (AJP 2024)', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Phương pháp hiện tại và tương lai trong điều trị rối loạn lo âu trẻ em. American J Psychiatry, 2024 (Q1, IF\u224818,0). 12 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Phương pháp', 'Hiệu quả', 'Ghi chú'],
    [['CBT (liệu pháp nhận thức HV)', 'Phục hồi 47\u201366%, đáp ứng 57\u201360%', 'Bằng chứng mạnh nhất'],
     ['SSRI (thuốc chống trầm cảm)', 'Hiệu quả vừa phải', 'Kết hợp CBT tốt hơn'],
     ['CBT + SSRI', 'Tốt hơn đơn trị', 'Chuẩn vàng hiện tại'],
     ['Can thiệp kỹ thuật số', 'Triển vọng, cần thêm NC', 'App, online CBT'],
     ['Kích thích thần kinh', 'Thử nghiệm, chưa đủ bằng chứng', 'TMS, tDCS']],
    widths=[4.0, 4.0, 4.0])
d.add_paragraph()
bl(d, 'So sánh: Zhameden 2025 (CBT 3/4 trầm cảm, 1/4 lo âu ở LMIC). 0 RCT từ Việt Nam.')
rh2(d, 'Phản biện')
bl(d, 'AJP Q1 IF=18 — tạp chí tâm thần hàng đầu. Hạn chế: tổng quan (review), chủ yếu bối cảnh phương Tây.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 AJP Q1 IF=18, tổng quan toàn diện cập nhật nhất.')
d.save('QT28_AJP_Treatment_2024.docx'); sys.stderr.write('QT28 OK\n')

# QT29 CBT Network Meta
d = make_doc()
bl(d, 'Tóm tắt: CBT interventions network meta-analysis (BMC Psychiatry 2025)', bold=True)
rb(d, 'Tên công trình')
bl(d, 'Hiệu quả các loại can thiệp khác nhau cho rối loạn lo âu ở trẻ em/VTN: Tổng quan hệ thống và phân tích tổng hợp mạng Bayesian. BMC Psychiatry, 2025 (Q1, OA). 30 RCTs, 1.711 trẻ. 14 trang.')
rb(d, 'Kết quả chính')
tbl(d, ['Can thiệp', 'Số RCTs', 'Hiệu quả', 'Xếp hạng'],
    [['CBT cá nhân', '\u2014', 'Hiệu quả nhất', '1'],
     ['CBT nhóm', '\u2014', 'Hiệu quả tốt', '2\u20133'],
     ['Thuốc (SSRI)', '\u2014', 'Hiệu quả', '3\u20134'],
     ['CBT + thuốc', '\u2014', 'Tốt nhưng ít RCT', '\u2014'],
     ['Tâm lý giáo dục', '\u2014', 'Hiệu quả thấp hơn', '5+'],
     ['Tổng', '30 RCTs', '1.711 trẻ', '\u2014']],
    widths=[3.5, 2.0, 3.5, 2.5])
d.add_paragraph()
bl(d, 'So sánh: Zhameden 2025 (6 RCTs LMIC, GRADE rất thấp, 0 RCT VN). AJP 2024 (CBT 47\u201366% phục hồi).')
rh2(d, 'Phản biện')
bl(d, 'BMC Q1 OA, 30 RCTs, network meta Bayesian — phương pháp nâng cao nhất. Hạn chế: cần đọc xếp hạng chi tiết. 0 RCT VN.')
rate(d, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50 BMC Q1, 30 RCTs, Bayesian network meta nâng cao.')
d.save('QT29_CBT_NetworkMeta_2025.docx'); sys.stderr.write('QT29 OK\n')
