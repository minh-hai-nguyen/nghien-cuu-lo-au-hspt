# -*- coding: utf-8 -*-
"""
Sách mỏng: Những mẩu chuyện về Madame Thảo
Dựa trên transcript cuộc họp + báo chí
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.6
for s in doc.sections:
    s.top_margin = Cm(3)
    s.bottom_margin = Cm(3)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2.5)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0,0,0)

def p(text, italic=False, size=13):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.italic = italic

def quote(text):
    """Trích dẫn lời Madame Thảo"""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True

def note(text):
    """Ghi chú nguồn"""
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(128, 128, 128)

# ===== BÌA =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('', level=0)
r = title.add_run('NHỮNG MẨU CHUYỆN\nVỀ MADAME THẢO')
r.font.name = 'Times New Roman'
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('TS. Nguyễn Thị Phương Thảo\nChủ tịch HĐQT Vietjet')
r.font.name = 'Times New Roman'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()
doc.add_paragraph()
epi = doc.add_paragraph()
epi.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = epi.add_run('"Mỗi con người đều có thể trở nên vĩ đại.\nLàm những việc lớn có ý nghĩa,\nmà từ những hoạt động rất là đơn giản, giản dị."')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.italic = True

doc.add_page_break()

# ===== LỜI MỞ ĐẦU =====
h('Lời mở đầu')
p('Quyển sách nhỏ này tập hợp những câu chuyện do chính TS. Nguyễn Thị Phương Thảo — Chủ tịch HĐQT Vietjet — kể lại trong một cuộc họp nội bộ về chương trình đào tạo lãnh đạo. Đây là lần đầu tiên những câu chuyện này được ghi nhận lại một cách có hệ thống.')
p('Từ cô sinh viên 17 tuổi một mình sang Nga, đến nữ tỷ phú tự thân đầu tiên của Việt Nam và Đông Nam Á — hành trình của Madame Thảo không chỉ là câu chuyện kinh doanh, mà là câu chuyện về những giá trị mà một con người có thể mang lại cho đất nước mình.')
p('"Tất cả những hoạt động của mình cho người dân, nó rất là thiết thực, ý nghĩa. Nó xuyên suốt mấy chục năm. Bây giờ mình phải lan tỏa tinh thần ấy cho nhân viên, cán bộ của mình, để mọi người tự hào, mọi người tự tin rằng mình sẽ làm được, bất chấp tất cả những khó khăn, thách thức từ bên ngoài."', italic=True)
doc.add_page_break()

# ===== CHƯƠNG 1: KHỞI ĐẦU =====
h('Chương 1: Cô gái 17 tuổi ở Moscow')
p('Năm 1987, cô gái Hà Nội Nguyễn Thị Phương Thảo, 17 tuổi, một mình lên đường sang Liên Xô du học. Cô theo học ngành Kinh tế Tài chính tại Cao đẳng Kinh tế Quốc dân Moscow và Học viện Thương mại Moscow.')
p('Ngay từ năm hai đại học, Phương Thảo đã bắt đầu buôn bán — phân phối máy fax, nhựa và cao su từ Nhật Bản, Hồng Kông và Hàn Quốc sang Liên Xô. Ở tuổi 21, cô đã trở thành triệu phú đô la.')
p('Nhưng câu chuyện không dừng ở việc kiếm tiền. Trong cuộc họp nội bộ, Madame Thảo kể lại hành trình tiếp theo:')
quote('"Việt Nam mình mới thoát ra khỏi cái đất nước là con nợ, nợ xấu ấy, bắt đầu đi ra thế giới. Thì mình lại đi tư vấn phát hành những cái trái phiếu chính phủ đầu tiên ở thị trường chứng khoán New York. Thành công những cái tỷ đô la đầu tiên. Rồi thành lập những cái ngân hàng thương mại đầu tiên."')
p('Bà đã góp vốn thành lập Ngân hàng Techcombank, tiếp đó là Ngân hàng VIB — những ngân hàng thương mại đầu tiên của Việt Nam trong thời kỳ đổi mới.')
note('Nguồn: Transcript cuộc họp nội bộ + bazaarvietnam.vn + interstellas.com.vn')
doc.add_page_break()

# ===== CHƯƠNG 2: GIẤC MƠ BAY =====
h('Chương 2: Giấc mơ bay cho mọi người')
p('Năm 2007, Nguyễn Thị Phương Thảo sáng lập Vietjet Air — hãng hàng không tư nhân đầu tiên của Việt Nam với mô hình giá rẻ. Đến năm 2019, Vietjet đã chiếm 41,2% thị phần hàng không nội địa.')
p('Nhưng ít ai biết rằng quyết định chuyển hướng sang mô hình giá rẻ đến từ một cuộc gặp gỡ rất đời thường:')
quote('Một bà mẹ ở vùng cao hỏi: "Bao nhiêu tấn thóc để mua một vé máy bay?"')
p('Câu hỏi đó đã thay đổi toàn bộ định hướng kinh doanh của Vietjet — từ mô hình 5 sao sang hãng hàng không giá rẻ, để hàng triệu người Việt Nam lần đầu tiên có thể bay.')
p('Khi IPO trên sàn chứng khoán tháng 2/2017, Nguyễn Thị Phương Thảo trở thành nữ tỷ phú tự thân đầu tiên của Việt Nam, tỷ phú thứ hai của đất nước (sau Phạm Nhật Vượng), và là nữ tỷ phú duy nhất của Đông Nam Á.')
note('Nguồn: tuoitre.vn + Forbes + Wikipedia')
doc.add_page_break()

# ===== CHƯƠNG 3: NÂNG SÂN BAY =====
h('Chương 3: Nâng đôi cánh cho Tân Sơn Nhất')
p('Madame Thảo kể lại trong cuộc họp nội bộ:')
quote('"Năm 2013 mình làm, 2014 bắt đầu chạy. Dự án nâng cao năng lực điều hành cất hạ cánh ở sân bay Tân Sơn Nhất — nâng lên gấp đôi từ 30 lượt đến 60 lượt. Bây giờ đang vận hành 48 lượt. Từ 24 triệu lượt khách lên 42 triệu lượt khách."')
p('Dự án không chỉ phục vụ Vietjet. Toàn bộ ngành hàng không được hưởng lợi. ACV (Tổng công ty Cảng hàng không Việt Nam) trong mười mấy năm liền, mỗi năm đón thêm 16 triệu lượt khách, lợi nhuận hàng trăm triệu đô la mỗi năm.')
quote('"Chứ không phải là doanh số đâu. Thì làm sao mà chẳng có tiền đi xây sân bay."')
p('Đến năm 2025, sân bay Tân Sơn Nhất vận hành tới 48 lượt cất hạ cánh mỗi giờ vào cao điểm. Nhà ga T3 mới với công suất 20 triệu khách mỗi năm nâng tổng công suất lên 50 triệu khách mỗi năm.')
note('Nguồn: lời kể từ cuộc họp nội bộ + plo.vn, dantri.com.vn (2025). Con số "30 lên 60 lượt" là từ lời kể, báo chí xác nhận 48 lượt/giờ cao điểm.')
doc.add_page_break()

# ===== CHƯƠNG 4: GIẢI CỨU =====
h('Chương 4: Giải cứu sàn chứng khoán')
p('Khi sàn chứng khoán Việt Nam bị treo hệ thống — một cuộc khủng hoảng ảnh hưởng đến hàng trăm tỷ đô la và 4 triệu khách hàng — Madame Thảo đã tham gia giải cứu.')
quote('"Cả nước Mỹ có sàn New York, Nasdaq — 2 cái sàn chứng khoán. Cả nước Anh có London Stock Exchange. Mà nó bị treo thì giải cứu mấy trăm tỷ đô la ở đấy. Có hơn 100 công ty chứng khoán, 4 triệu khách hàng. Mỗi ngày giao dịch khoảng 1 tỷ đô la, mà cao là 2 tỷ đô la. Mình cũng nhảy vào giải cứu."')
p('Theo báo chí ghi nhận, đầu năm 2021, sàn HoSE (Sở Giao dịch Chứng khoán TP.HCM) bị nghẽn lệnh nghiêm trọng — hệ thống giao dịch do Hàn Quốc cung cấp không đáp ứng được khối lượng. Madame Thảo đã kết nối Sovico với FPT để xây dựng hệ thống giao dịch thay thế. Chỉ trong khoảng 100 ngày, hệ thống mới chính thức vận hành ngày 5/7/2021, chi phí khoảng 60 tỷ đồng — giải cứu thành công thị trường vốn cho cả nước.')
p('Chính phủ giữ được sàn. Các công ty, nhà đầu tư được hưởng lợi. Các công ty niêm yết có thị trường vốn để huy động và phát triển.')
quote('"Thế còn mình thì không cần cái gì cả."')
note('Nguồn: Transcript + tinnhanhchungkhoan.vn + tuoitre.vn (5/7/2021) + cafef.vn')
doc.add_page_break()

# ===== CHƯƠNG 5: NGOẠI GIAO =====
h('Chương 5: Ngoại giao trên bầu trời')
p('Mỗi hợp đồng mua máy bay là một cơ hội xây dựng quan hệ ngoại giao kinh tế. Madame Thảo giải thích:')
quote('"Mua máy bay Boeing là đi với nước Mỹ — ông Trump, ông Obama, ông Bush. Mua máy bay Airbus là đi với Pháp, Đức, Tây Ban Nha. Mua máy bay COMAC của Trung Quốc là chương trình của ông Tập Cận Bình. Là mình tạo nên xây dựng mối quan hệ ngoại giao, xây dựng kinh tế giữa các quốc gia, dân tộc. Mình sẽ giảm bớt đi những chiến tranh, những sự xung đột."')
p('Và không chỉ là hợp đồng:')
quote('"Máy bay mình bay đến đâu, biên giới đất nước nó mở ra. Mình mang cờ cho khắp thế giới, các đường bay mình đến."')
p('Ngày 15 tháng 4 năm 2021, Đại sứ Pháp tại Việt Nam Nicolas Warnery đã trao Huân chương Bắc đẩu Bội tinh (Ordre national de la Légion d\'honneur) cho Madame Thảo — huân chương cao quý nhất của nước Pháp, được Napoléon Bonaparte thiết lập năm 1802. Bà là nữ doanh nhân Việt Nam đầu tiên nhận vinh dự này, vì những đóng góp trong thúc đẩy quan hệ Việt-Pháp.')
p('Bà cũng được Đại học Harvard đưa vào tài liệu nghiên cứu và giảng dạy, được Business Insider Australia bầu chọn là một trong 100 người thay đổi kinh tế châu Á, và được tạp chí Tatler vinh danh là người có ảnh hưởng nhất châu Á trong lĩnh vực từ thiện.')
note('Nguồn: Transcript + baochinhphu.vn + vnexpress.net + tuoitre.vn (15/4/2021)')
doc.add_page_break()

# ===== CHƯƠNG 6: APEC =====
h('Chương 6: 90 ngày — khi Hàn Quốc mất 17 năm')
p('Vietjet tham gia phục vụ APEC 2 lần, đặc biệt năm 2017 — "Đưa Việt Nam vào thế giới và đưa thế giới về Việt Nam."')
p('Madame Thảo kể lại một so sánh ấn tượng:')
quote('"Cái giải pháp chiếu sáng mình làm có 90 ngày thôi. Bình thường nó phải mất 12 tháng mới xong. Trong khi đó Hàn Quốc triển khai 12 năm mà không chạy, đến lúc bàn giao đã 17 năm. Tốn 100 triệu đô rồi, không bàn giao được."')
quote('"Xây cung hội nghị 4 tháng."')
quote('"Như là cái tòa Marina này, nó đã mang đến một cái sinh khí cho cả cái thành phố này."')
p('Còn khi tình hình quốc tế căng thẳng:')
quote('"Ngày 9 tháng 1 năm ngoái, không ai mời mình đâu, nhưng mà thấy thời cuộc nó căng thẳng. Mình có mối quan hệ tốt với ông Trump, Đại sứ quán Mỹ mời mình, thì mình bay luôn, mang máy bay đi luôn. Cũng tạo nên thiện cảm, cũng đỡ căng thẳng."')
note('Nguồn: lời kể từ cuộc họp nội bộ. APEC 2017 tại Đà Nẵng xác nhận qua báo chí (vnexpress.net, baochinhphu.vn). Chi tiết "90 ngày" và "Hàn Quốc 17 năm" là từ lời kể, cần xác nhận thêm.')
doc.add_page_break()

# ===== CHƯƠNG 7: COVID =====
h('Chương 7: Mùa COVID — lao đi vì người dân')
p('Khi đại dịch COVID-19 bùng phát, Vietjet và Madame Thảo đã hành động ngay lập tức:')
quote('"Mùa Covid mình lao đi chở vắc-xin, thiết bị y tế, bao nhiêu những bao gạo. Không có những cái đó thì làm sao mà người ta được yên lòng."')
p('Theo báo chí ghi nhận, Vietjet đã chuyên chở miễn phí hơn 2.000 bác sĩ và nhân viên y tế từ miền Bắc và miền Trung vào TP.HCM. Hàng trăm máy thở và thiết bị y tế từ Đức được vận chuyển miễn phí về Việt Nam. Vietjet và HDBank tặng 10 xe cứu thương, tài trợ hơn 1.000 giường bệnh.')
p('Chương trình "Bữa cơm yêu thương" do Madame Thảo khởi xướng đã cung cấp hơn 600.000 suất ăn trong hơn một tháng. Tổng cộng ước tính hơn 70 tỷ đồng cho công tác thiện nguyện và phòng chống dịch tại TP.HCM.')
note('Nguồn: Transcript + vietnamplus.vn + thanhnien.vn + hdbank.com.vn')
doc.add_page_break()

# ===== CHƯƠNG 8: THIẾT KẾ =====
h('Chương 8: Cái thang máy dành cho quan tài')
p('Trong cuộc họp, Madame Thảo chia sẻ triết lý thiết kế mà ít ai biết:')
quote('"Tất cả những cái gì mà mình làm — thiết kế một cái nhà chung cư, là phải có cái thang máy lớn để dành cho cáng cứu thương và quan tài. Từ cái thiết kế là mình cũng nghĩ đến con người — lúc người ta bệnh, người ta cần phải đưa đi cáng. Một cái thang máy ở nhà chung cư mà không thiết kế ra như thế."')
quote('"Ở đâu có bậc thang là mình phải có cái đường dốc cho xe nôi, cho xe người khuyết tật. Đấy là tinh thần mình đưa vào mọi hoạt động."')
p('Đó là cách Madame Thảo nghĩ về con người — không phải ở những lúc vui vẻ, mà ở những lúc con người cần được giúp đỡ nhất: khi bệnh, khi già, khi khuyết tật. Và tinh thần đó không chỉ nằm trong thiết kế bất động sản, mà trong mọi hoạt động của bà.')
note('Nguồn: Transcript cuộc họp nội bộ')
doc.add_page_break()

# ===== CHƯƠNG 9: XÃ HỘI =====
h('Chương 9: 10.000 đôi mắt mỗi năm')
p('Madame Thảo kể về hoạt động xã hội xuyên suốt mấy chục năm — từ cờ vua, bóng đá, đến y tế:')
quote('"Cờ vua Việt Nam lên top 7 thế giới. Làm bóng đá futsal. Rồi bóng đá nữ, rồi thể thao người khuyết tật."')
quote('"Các chương trình mổ mắt, mổ tim, khám bệnh, phát thuốc. Phát bảo hiểm y tế."')
p('Một thành viên trong cuộc họp bổ sung:')
quote('"Riêng cái chương trình mổ mắt cho người nghèo — một năm 10.000 ca mổ mắt. Rồi tổ chức khám chữa bệnh ở khắp nơi."')
p('Madame Thảo cũng khởi xướng những phong trào mới lạ:')
quote('"Kêu gọi cho bà con bắt đầu đưa hàng hóa lên trên TikTok để bán ấy. Thực ra mình là những cái người đi phát động."')
p('Giải cờ vua quốc tế HDBank, dự án "Tôi yêu Tổ quốc tôi" với hơn 6 triệu đoàn viên thanh niên tham gia, các chương trình thiện nguyện tại Philippines, Ấn Độ — tất cả đều nằm trong hệ sinh thái đóng góp xã hội mà Madame Thảo dẫn dắt.')
note('Nguồn: Transcript + tuoitre.vn + thanhnien.vn')
doc.add_page_break()

# ===== CHƯƠNG 10: TINH THẦN =====
h('Chương 10: Cái hồn của đào tạo')
p('Trong cuộc họp, Madame Thảo trăn trở về tinh thần nội bộ:')
quote('"Đấy là cái tinh thần mà mình đưa vào cho cái đào tạo của mình. Nó khác với những cái hoạt động thông thường mà sau này mình cứ làm cho nó mai một đi, làm cho nó đi đâu mất. Tức là nó thiếu cái kiểu giống như con người, nó thiếu cái hồn của nó."')
p('Bà kể lại hình ảnh tổ bay ngày xưa:')
quote('"Ngày xưa tổ bay là briefing ở tại công ty, thậm chí là còn xếp hàng ngay ngắn. Đi một lượt, kiểm tra lại với nhau, xong rồi diễu ở trong văn phòng ấy. Nhân viên đầu tiên nhìn cái tổ bay này chuẩn bị làm nhiệm vụ, thì người ta tự hào, người ta thấy cái giá trị, cái ý nghĩa của công việc."')
p('Và thực tế đáng lo:')
quote('"Chứ còn bây giờ là cứ đầu tắt mặt tối, ra khỏi nhà 2-3 giờ sáng, ra sân bay rồi về đi ngủ. Kiếm một số tiền thì không nhỏ, nhưng mà cũng đủ sống. Không có thời gian tiêu. Người ta rời đi không phải vì chế độ của mình. Chế độ của mình là chế độ tốt. Nhưng mà cái môi trường, người ta cảm thấy là người ta không có ý nghĩa."')
note('Nguồn: Transcript cuộc họp nội bộ')
doc.add_page_break()

# ===== CHƯƠNG 11: NHÂN QUẢ =====
h('Chương 11: Nhân quả')
p('Triết lý quản lý của Madame Thảo đơn giản nhưng sâu sắc:')
quote('"Mình tin vào nhân quả. Nhân nào quả đấy. Mình làm chưa tốt thì cuộc sống sẽ có cái quả tương ứng. Chứ không cần phải có hành động gì xử lý căng thẳng."')
quote('"Cuộc sống, người ta là con người, người ta có quyền sai. Thì mình hãy tạo cho người ta cơ hội để có thể khắc phục."')
p('Về công bằng:')
quote('"Tạo nên sự công bằng trong tổ chức là cần thời gian, cần cơ chế. Chứ không phải bây giờ cấp dưới đụng chạm một cái gì đấy mà mình cho rằng là sai, thì mình xử lý rất là mạnh. Trong khi đó lãnh đạo ở trên không gương mẫu, thì đấy nó không công bằng."')
p('Về coaching:')
quote('"Mỗi người phải trở thành người coach cho những người xung quanh. Cũng phải biết cách học hỏi."')
p('Đó là triết lý của người phụ nữ từng nói: "Tiền nhiều để hiện thực hóa những ước mơ cao đẹp và giúp đỡ được nhiều người hơn."')
note('Nguồn: Transcript + tuoitre.vn')
doc.add_page_break()

# ===== KẾT =====
h('Lời kết')
p('Những câu chuyện trong cuốn sách nhỏ này chỉ là một phần rất nhỏ trong hành trình mấy chục năm của Madame Thảo. Như bà nói, đây là lần đầu tiên những câu chuyện này được thống kê lại.')
p('Từ cô sinh viên ở Moscow đến nữ tỷ phú Forbes, từ giải cứu sàn chứng khoán đến 10.000 ca mổ mắt mỗi năm, từ ngoại giao trên bầu trời đến cái thang máy dành cho quan tài — tất cả đều hướng về một thông điệp:')
quote('"Mỗi con người đều có thể trở nên vĩ đại. Làm những việc lớn có ý nghĩa, mà từ những hoạt động rất là đơn giản, giản dị của mình thôi."')
doc.add_paragraph()
p('Sách được biên soạn từ bản ghi cuộc họp nội bộ, có tham khảo thêm nguồn báo chí. Một số chi tiết cần được xác minh từ nguồn chính thức trước khi sử dụng trong tài liệu đối ngoại.', size=10, italic=True)

fname = 'Sách - Những mẩu chuyện về Madame Thảo.docx'
doc.save(fname)
print('Done: ' + fname)
