# -*- coding: utf-8 -*-
"""Dịch A5 Dong 2025 PLOS ONE — DASS-21 secondary school students Ya'an, TQ"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                t += ' ' + cell.text
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

# ========== A5 Dong 2025 PLOS ONE ==========
print('A5 Dong 2025 PLOS (full)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1371/journal.pone.0328785', size=10)
add_heading(doc, 'Tỷ lệ và các yếu tố quyết định trầm cảm, lo âu và căng thẳng ở học sinh trung học', 1)
h = doc.add_paragraph()
r = h.add_run('Prevalence and determinants of depression, anxiety, and stress among secondary school students')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'Tingting Dong (1), Yumei Wang (2), Yuanjie Lin (3)'),
    ('Cơ quan', '(1) BV Văn phòng Chính phủ Tây Tạng tại Thành Đô; (2) ĐH Tây Hoa, Khoa KH Vật liệu; (3) Trường YTCC, ĐH Tứ Xuyên — Thành Đô, TQ'),
    ('Tạp chí', 'PLOS ONE (Q1, IF ≈ 3,7)'),
    ('Xuất bản', '2025-09-03, Vol. 20(9): e0328785, 10 trang'),
    ('DOI', '10.1371/journal.pone.0328785'),
    ('Loại NC', 'Cắt ngang — khảo sát + hồi quy logistic nhị phân'),
    ('Mẫu', '2.716 HS trung học Ya An, Tứ Xuyên TQ — 1.230 nam (45,3%) + 1.486 nữ (54,7%)'),
    ('Công cụ', 'DASS-21 (Depression-Anxiety-Stress Scale 21 mục) + bảng câu hỏi nhân khẩu học'),
    ('Thời gian', '03-06/2022 — 4 trường THCS+THPT Ya An'),
    ('Kinh phí', 'Hiệp hội Y học Thành Đô (Project No. 2024350)'),
    ('Truy cập', 'Open Access — Creative Commons BY'),
])
add_page_ref(doc, '1–10', 'PLOS ONE', 'Vol. 20(9), e0328785, 2025')

add_heading(doc, 'BẢNG VIẾT TẮT', 2)
add_abbreviation_table(doc, [
    ('DASS-21', 'Depression-Anxiety-Stress Scale 21 mục — thang đo trầm cảm/lo âu/căng thẳng'),
    ('OR', 'Odds Ratio — tỷ số chênh'),
    ('CI', 'Confidence Interval — khoảng tin cậy 95%'),
    ('VIF', 'Variance Inflation Factor — hệ số phóng đại phương sai'),
    ('JHSE', 'Junior High School Examination — kỳ thi vào THPT TQ'),
    ('Gao Kao', 'Kỳ thi vào ĐH TQ — National Unified Examination'),
    ('GBD', 'Global Burden of Disease — gánh nặng bệnh tật toàn cầu'),
    ('THCS', 'Trung học cơ sở (junior high — lớp 7-9)'),
    ('THPT', 'Trung học phổ thông (senior high — lớp 10-12)'),
])

add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Bối cảnh: Trẻ vị thành niên (10–19 tuổi theo WHO) là cửa sổ phát triển SKTT then chốt. Tại TQ, 68% rối loạn tâm thần người lớn khởi phát trước 18 tuổi. Hệ thống giáo dục TQ với JHSE và Gao Kao tạo áp lực lớn. NC nhằm khảo sát tỷ lệ + yếu tố ảnh hưởng trầm cảm, lo âu, căng thẳng ở HS trung học để cung cấp cơ sở cho NC + can thiệp tiếp theo.')
add_p(doc, 'Phương pháp: Cắt ngang, khảo sát điện tử (QuestionStar) tháng 3-6/2022 ở 4 trường trung học Ya An, Tứ Xuyên. Sampling cụm. Mẫu mục tiêu n=1.844 (DE=4, attrition 20%). Thu được 2.716/2.900 phiếu hợp lệ (93,66%). DASS-21 + bảng nhân khẩu học. Hồi quy logistic nhị phân (Model 1 đơn biến → Model 2 đa biến loại p>0,1 hoặc VIF>10).')

p = doc.add_paragraph()
r = p.add_run('Kết quả: 24,4% có triệu chứng trầm cảm; 41,4% triệu chứng lo âu; 15,6% căng thẳng. Tỷ lệ lo âu (41,4%) CAO HƠN trầm cảm và stress rõ rệt. Hồi quy đa biến cho thấy: học THPT, xếp hạng học tập <60%, ÍT/KHÔNG hỗ trợ tâm lý từ gia đình + bạn bè là các yếu tố chính.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_p(doc, 'Kết luận: HS THPT, HS xếp hạng học tập <60%, HS thiếu hỗ trợ gia đình/bạn bè là nhóm NGUY CƠ CAO nhất. Cần chiến lược can thiệp PHÂN TẦNG theo yếu tố nguy cơ cụ thể.')

add_heading(doc, 'GIỚI THIỆU', 2)
add_p(doc, 'WHO định nghĩa VTN là 10–19 tuổi — cửa sổ phát triển SKTT quan trọng. Ở TQ, THCS gồm lớp 7–9 (12–15t) và THPT lớp 10–12 (16–18t). Hai kỳ thi lớn JHSE (vào THPT) và Gao Kao (vào ĐH) tạo áp lực khác nhau. Bằng chứng gần đây: yếu tố nguy cơ chính (áp lực HT, quan hệ gia đình, cạnh tranh bạn bè) cố kết từ đầu VTN. NC dọc cho thấy 68% rối loạn tâm thần người lớn khởi phát trước 18 tuổi.')
add_p(doc, 'GBD 2019: trầm cảm chiếm phần lớn nhất (37,3%), tiếp theo là lo âu (22,9%) — gánh nặng nặng nề lên xã hội + gia đình.')
add_p(doc, 'Mojtabai 2016 (Mỹ): trầm cảm nặng VTN tăng từ 8,7% → 11,3% (2005-2014). Indonesia 1.381 HS: 60% có triệu chứng trầm cảm. TQ: trầm cảm THPT ~27,12% và đang tăng. NC này nhằm hiểu tình trạng trầm cảm/lo âu/stress + phân tích yếu tố ảnh hưởng.')

add_heading(doc, 'PHƯƠNG PHÁP', 2)
add_p(doc, 'Bối cảnh: Cắt ngang ở Ya An, Tứ Xuyên TQ. Đã được phê duyệt đạo đức (Số 50, 2022 — UBĐĐ BV Văn phòng Chính phủ Tây Tạng). Đồng thuận có chữ ký từ HS và phụ huynh.')
add_p(doc, 'Tiêu chí chọn: (1) HS đang học; (2) đồng thuận kép HS + phụ huynh; (3) hoàn thành phiếu điện tử độc lập.')
add_p(doc, 'Tính cỡ mẫu: P=0,5, d=0,05, α=0,05 → Z=1,96 → n=384. Sampling cụm DE=4 → 1.536. Cộng 20% attrition → mục tiêu 1.844. Thực tế đạt 2.716 hợp lệ.')
add_p(doc, 'Công cụ DASS-21 (Lovibond 1995): 21 mục, mỗi mục 0–3. Cắt ngưỡng: Trầm cảm ≥10 (10–13 nhẹ, 14–20 TB, ≥21 nặng/rất nặng); Lo âu ≥8 (8–9 nhẹ, 10–14 TB, ≥15 nặng); Stress ≥15 (15–18 nhẹ, 19–25 TB, ≥26 nặng). DASS-21 đã validate trên dân số TQ với độ tin cậy + tính giá trị tốt.')
add_p(doc, 'Phân tích: Mô tả (mean±SD, n+%). So sánh nhóm: t-test/Wilcoxon (liên tục), chi-square (phân loại). Hồi quy logistic: Model 1 đơn biến; Model 2 loại p>0,1 hoặc VIF>10. Ngưỡng p<0,05 hai phía. Phần mềm R.')

add_heading(doc, 'KẾT QUẢ', 2)
add_heading(doc, 'Bảng 1. Phân bố DASS-21 ở 2.716 HS trung học (Bảng 2 bài gốc)', 3)
add_table(doc,
    ['Triệu chứng', 'Bình thường', 'Nhẹ', 'Trung bình', 'Nặng', 'Rất nặng', 'Tổng %'],
    [['Trầm cảm', '2.053 (75,59%)', '337 (12,41%)', '230 (8,47%)', '48 (1,77%)', '48 (1,77%)', '24,41%'],
     ['Lo âu', '1.591 (58,58%)', '397 (14,62%)', '499 (18,37%)', '111 (4,09%)', '118 (4,34%)', '41,42%'],
     ['Stress', '2.293 (84,43%)', '233 (8,58%)', '114 (4,2%)', '58 (2,14%)', '18 (0,66%)', '15,57%']],
    widths=[2.5, 2.3, 1.8, 2.2, 1.8, 2.0, 1.8])
add_p(doc, 'Lo âu (41,4%) cao hơn rõ rệt trầm cảm (24,4%) và stress (15,6%). Trong 41,4% lo âu: 14,62% nhẹ, 18,37% TB, 4,09% nặng, 4,34% rất nặng — gần 9% lo âu nặng/rất nặng.', italic=True)

add_heading(doc, 'Bảng 2. Yếu tố ảnh hưởng — phân tích đơn biến (Bảng 1 bài gốc)', 3)
add_table(doc,
    ['Yếu tố', 'Nhóm', 'Trầm cảm có (%)', 'Lo âu có (%)', 'Stress có (%)', 'p (TC/LA/Stress)'],
    [['Tuổi (mean)', '—', '15,5±1,7', '15,4±1,8', '15,4±1,7', '<0,001/<0,001/0,003'],
     ['Giới', 'Nữ vs Nam', '55,4% vs 54,5%', '56,1% vs 53,7%', '54,1% vs 54,8%', '0,736/0,241/0,837'],
     ['Lớp', 'THPT vs THCS', '76,5% vs 23,5%', '73,8% vs 26,2%', '74,7% vs 25,3%', '<0,001/<0,001/0,016'],
     ['Học bán trú', 'Nội trú vs đi về', '69,5% vs 30,5%', '76% vs 24%', '70,2% vs 29,8%', '0,014/0,008/0,137'],
     ['Nơi ở', 'Nông thôn vs Thành thị', '59,3% vs 40,7%', '60,8% vs 39,2%', '57% vs 43%', '0,032/<0,001/0,57'],
     ['Xếp hạng <60%', 'Có vs nhóm khác', '30%', '23,4%', '28,1%', '<0,001/0,824/0,077'],
     ['Hỗ trợ gia đình HIẾM', 'Hiếm vs thường', '11,3% vs 38,6%', '8% vs 45,2%', '12,5% vs 41,4%', '<0,001/<0,001/<0,001'],
     ['Hỗ trợ bạn HIẾM', 'Hiếm vs thường', '8,9% vs 60%', '6,2% vs 62,4%', '9,9% vs 59,1%', '<0,001/<0,001/<0,001'],
     ['Tìm trợ giúp THỤ ĐỘNG', 'Thụ động vs chủ động', '17,2% vs 17,8%', '11,9% vs 21,6%', '18,7% vs 19,9%', '<0,001/<0,001/<0,001'],
     ['KHÔNG có người tâm sự', 'Không vs có gia đình', '21,9% vs 14,3%', '15,4% vs 18,3%', '23,2% vs 12,8%', '<0,001/<0,001/<0,001']],
    widths=[3.0, 2.8, 2.5, 2.5, 2.5, 2.7])
add_p(doc, 'GIỚI tính KHÔNG khác biệt (cả 3 chỉ số p>0,2). Trái với QT08 Wen 2020 (nữ > nam OR=11,58 lo âu) — có thể do bài này đo cả 3 trục hoặc do sample TQ Tứ Xuyên khác đặc thù.', italic=True)

add_heading(doc, 'Bảng 3. Hồi quy logistic Model 2 đa biến (rút gọn) — Bảng 3 bài gốc', 3)
add_table(doc,
    ['Yếu tố', 'Trầm cảm OR (95% CI)', 'Lo âu OR (95% CI)', 'Stress OR (95% CI)'],
    [['Tuổi', '1,24 (1,08–1,43)***', '2,06 (1,79–2,38)***', '1,08 (0,91–1,27)'],
     ['THPT vs THCS', '0,65 (0,37–1,13)', '0,10 (0,06–0,17)***', '0,92 (0,48–1,76)'],
     ['Học nội trú vs đi về', '0,78 (0,62–0,99)*', '1,23 (1,00–1,50)*', '0,88 (0,69–1,12)'],
     ['Xếp hạng <60% (so Top 10%)', '1,62 (1,17–2,25)**', '0,98 (0,70–1,36)', '1,22 (0,87–1,71)'],
     ['Hỗ trợ GĐ "More" (so seldom)', '0,45 (0,30–0,68)***', '0,69 (0,47–1,01)', '0,53 (0,35–0,82)**'],
     ['Hỗ trợ bạn "More" (so seldom)', '0,64 (0,41–1,00)*', '0,73 (0,47–1,12)', '0,57 (0,36–0,91)*'],
     ['Tìm giúp CHỦ ĐỘNG (so thụ động)', '0,48 (0,32–0,73)***', '0,70 (0,47–1,04)', '0,58 (0,37–0,92)*'],
     ['Tâm sự với gia đình (so không)', '0,22 (0,15–0,33)***', '0,27 (0,19–0,40)***', '0,23 (0,15–0,36)***'],
     ['Tâm sự với bạn (so không)', '0,37 (0,26–0,53)***', '0,42 (0,29–0,60)***', '0,48 (0,33–0,70)***']],
    widths=[4.5, 3.8, 3.8, 3.8])
add_p(doc, '*p<0,05, **p<0,01, ***p<0,001. Các yếu tố BẢO VỆ MẠNH NHẤT (OR thấp nhất): tâm sự với gia đình giảm trầm cảm 78% (OR=0,22), lo âu 73%, stress 77%. Tâm sự với bạn giảm 60%/58%/52%. Hỗ trợ gia đình "More" giảm trầm cảm 55%.', italic=True)
add_p(doc, 'Lưu ý kỹ thuật: OR=0,10 cho "THPT vs THCS" trong cột lo âu là KHÁC chiều với phân tích đơn biến (THPT có lo âu cao hơn) — điều này phản ánh sau khi kiểm soát TUỔI (OR=2,06 cho lo âu) đã hấp thụ phần lớn hiệu ứng "lên lớp", nên biến grade trở thành chống chiều. Cần cẩn trọng diễn giải đa cộng tuyến.', italic=True)

add_heading(doc, 'THẢO LUẬN', 2)
add_p(doc, 'Lo âu 41,4% > trầm cảm 24,4% > stress 15,6%. Hồi quy đa biến: học THPT, ở nội trú dài hạn, gia đình đơn thân, xếp hạng <60%, thiếu hỗ trợ GĐ + bạn, ít chủ động tìm giúp đỡ là các yếu tố nguy cơ chung của cả 3.')
add_p(doc, 'HS THPT có tỷ lệ cao hơn THCS — gốc rễ ở áp lực Gao Kao đa chiều: >10 triệu HS thi mỗi năm, cạnh tranh 6–8 triệu suất ĐH. Tỷ lệ enroll ĐH 45,7%–57,8% (2017-2021). HS THPT trung bình HỌC 11,2 GIỜ/NGÀY (tăng 31,8% so THCS) — kích hoạt liên tục trục HPA, tăng lo âu 32,7%. Tại Cát Lâm (tỷ lệ vào ĐH top 1 chỉ 20,3%), 61,7% HS đặt mục tiêu "Double First-Class" và 70,2% có lo âu lâm sàng do khoảng cách mục tiêu/thực tế. 82,4% phụ huynh có kỳ vọng học tập rõ ràng, 48,6% kỳ vọng vào ĐH top → 60,3% HS sợ "làm người khác thất vọng". Lịch 12 giờ/ngày (7h–21h) + 3,2 bài kiểm tra/tuần → 70,1% HS có triệu chứng cơ thể hóa rõ.')
add_p(doc, 'Gia đình đơn thân: nguy cơ cao hơn gia đình đầy đủ. Huang et al.: cha mẹ + ông bà có vai trò bảo vệ. Thiếu hỗ trợ ổn định + lo định kiến xã hội → cô lập + tăng gánh nặng tâm lý. VTN ở giai đoạn nhận thức bản thân + điều tiết cảm xúc chưa trưởng thành — dễ tổn thương trước yếu tố ngoại cảnh.')
add_p(doc, 'Hành vi tìm giúp: HS thường THIẾU kênh tìm trợ giúp hiệu quả. Mô hình gia đình độc đoán + thiếu cởi mở → giao tiếp kém. Cạnh tranh bạn bè → khó tìm hỗ trợ peer. Tự trọng cao + nhận thức "tìm giúp = yếu" → chịu đựng một mình. Cần xây dựng hệ thống hỗ trợ ĐA TẦNG: gia đình + trường + xã hội.')

add_heading(doc, 'HẠN CHẾ', 2)
add_p(doc, '(1) Cắt ngang — chỉ là snapshot, cần NC dọc xác nhận quan hệ nhân quả. (2) Số HS THCS ít hơn THPT do nhiều HS THCS không có điện thoại riêng. (3) Không thu thập tiền sử gia đình rối loạn cảm xúc → có thể lệch.')
add_p(doc, 'Kết luận: SKTT HS trung học gắn chặt với LỚP HỌC, THÀNH TÍCH HỌC TẬP, HỖ TRỢ TÂM LÝ TỪ GIA ĐÌNH + BẠN BÈ.')

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'PLOS ONE Q1, Open Access — peer review history công khai.',
    'Mẫu LỚN n=2.716 — đủ power cho hồi quy đa biến nhiều biến.',
    'Đo CẢ 3 trục DASS-21 (trầm cảm + lo âu + stress) — toàn diện hơn các NC đo riêng.',
    'Phân tích 2 model (đơn + đa biến với kiểm tra VIF) — giảm nguy cơ confounding.',
    'Phát hiện CƠ CHẾ BẢO VỆ rõ ràng: tâm sự với gia đình OR=0,22 (giảm trầm cảm 78%) — mạnh hơn cả "More family support" — gợi ý rằng KÊNH GIAO TIẾP quan trọng hơn mức độ hỗ trợ chung.',
    'Phù hợp đề tài VN: Trần Thảo Vi VN21 (β=4,73 học thêm), Wen QT08 (OR=11,58 áp lực HT), UNICEF VN22 (47% học thêm >3h/tuần). Xác nhận áp lực HT là cơ chế chung TQ + VN.',
    'Đặc biệt: Lo âu 41,4% — tương đương Hải Phòng (39,3% THPT) và Long An VN (57,2%) → tỷ lệ ~40-50% có thể là baseline VTN châu Á.',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Hạn chế:', bold=True)
for s in [
    'CẮT NGANG — không thể kết luận nhân quả. Cần NC dọc (như VN21 Trần Thảo Vi 3 năm).',
    'Convenience sampling 4 trường ở 1 thành phố — không đại diện toàn TQ. Ya An là vùng dân tộc thiểu số Tứ Xuyên, kết quả có thể khác đô thị Bắc Kinh/Thượng Hải.',
    'GIỚI tính p>0,2 — TRÁI với hầu hết NC khác (Wen 2020 OR=11,58 nữ, Hoa 2024 nữ > nam). Cần kiểm tra bias khảo sát điện tử (nữ trả lời nhiều hơn).',
    'Tự báo cáo qua điện thoại + WeChat → có thể có social desirability bias.',
    'Tỷ lệ lo âu 41,4% > trầm cảm 24,4% — bất thường so nhiều NC khác (thường ngược) → cần kiểm tra việc dùng cut-off DASS-21 ≥8 cho lo âu (thấp).',
    'KHÔNG báo cáo tỷ lệ phản hồi HS riêng biệt giữa 4 trường → khó đánh giá selection bias.',
    'Bảng 3 có vấn đề diễn giải: OR=0,10 THPT cho lo âu sau khi kiểm soát tuổi trở nên KHÓ HIỂU — đa cộng tuyến với "Tuổi" (OR=2,06).',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Gap & gợi ý cho VN:', bold=True)
for s in [
    'PHÁT HIỆN MỚI: Tâm sự với gia đình giảm 78% nguy cơ trầm cảm — mạnh hơn mức hỗ trợ chung. Gợi ý CAN THIỆP TRƯỜNG VN: tập huấn KỸ NĂNG GIAO TIẾP cha-con, không chỉ tăng "hỗ trợ" chung chung.',
    'Xếp hạng <60% là yếu tố độc lập (OR=1,62 trầm cảm) — phù hợp với UNICEF VN22 + Wen QT08. Gợi ý: cần cơ chế hỗ trợ HS xếp hạng thấp KHÔNG dán nhãn.',
    'Cần NC VN dùng DASS-21 đo CẢ 3 trục ở mẫu lớn (>2.000) THPT để so sánh trực tiếp.',
    'NC dọc 3 năm (như VN21) đo cả 3 trục DASS-21 sẽ là đóng góp quan trọng.',
]:
    add_red(doc, f'• {s}')

add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. PLOS ONE Q1, n=2.716, đo cả 3 trục DASS-21, hồi quy đa biến, phát hiện cơ chế bảo vệ rõ ràng. Hạn chế chính: cắt ngang + 1 thành phố. Rất phù hợp đối chiếu TQ↔VN cho yếu tố áp lực HT + hỗ trợ xã hội.', bold=True)

save(doc, '54_Dong_PLOS_DASS_2025.docx')
print('DONE A5')
