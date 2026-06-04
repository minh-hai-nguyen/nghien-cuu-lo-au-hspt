# -*- coding: utf-8 -*-
"""
Sửa 8 lỗi trích dẫn cuối (phát hiện qua re-audit 6 agent 16/05/2026):
- Bài 1: Lee (đã có tập/số 27(2),391-406,2026), UNICEF (tác giả thật =
  Viện Xã hội học và cs), Xu (4 đồng tác giả cuối sai), Brunborg (sai thứ tự).
- Bài 2: Li (danh sách tác giả + số bài e13), UNICEF.
- Bài 03: Copeland (thêm phụ đề), Essau (tách họ tác giả 4: Ho, M.-H. R.).
Patch trực tiếp vào v6 / v4.
"""
import os
from docx import Document

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn"
B1 = os.path.join(ROOT, "Bai1_YTNC_HSTHCS_v6_16052026.docx")
B2 = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v6_16052026.docx")
B03 = os.path.join(ROOT, "bài-03", "Công Thị Hằng_v4_đã sửa.docx")

UNICEF_NEW = ("Viện Xã hội học, Đại học Queensland, & Đại học Johns Hopkins. (2022). "
              "Viet Nam adolescent mental health survey (V-NAMHS): Report on main "
              "findings. Viện Xã hội học.")


def setp(p, t):
    if p.runs:
        p.runs[0].text = t
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(t)


def whole(doc, sub, new):
    """Thay TOÀN BỘ đoạn chứa sub."""
    for p in doc.paragraphs:
        if sub in p.text:
            setp(p, new)
            return True
    return False


def part(doc, sub, new):
    """Thay cụm con trong mọi đoạn chứa sub."""
    n = 0
    for p in doc.paragraphs:
        if sub in p.text:
            setp(p, p.text.replace(sub, new))
            n += 1
    return n


# ===== BÀI 1 =====
d1 = Document(B1)
r = []
r.append(("TLTK Lee", whole(d1, "Lee, J., Choo, H., Zhang, Y., Cheung, H. S., Zhang, Q.",
    "Lee, J., Choo, H., Zhang, Y., Cheung, H. S., Zhang, Q., & Ang, R. P. (2026). "
    "Cyberbullying victimization and mental health symptoms among children and "
    "adolescents: A meta-analysis of longitudinal studies. Trauma, Violence, & "
    "Abuse, 27(2), 391–406. https://doi.org/10.1177/15248380241313051")))
r.append(("TLTK UNICEF", whole(d1, "UNICEF Vietnam. (2022)", UNICEF_NEW)))
r.append(("TLTK Xu", whole(d1, "Xu, Q., Mao, Z., Wei, D., Liu, P., Fan, K.",
    "Xu, Q., Mao, Z., Wei, D., Liu, P., Fan, K., Wang, J., Wang, X., Lou, X., Lin, H., "
    "Wang, C., & Wu, C. (2021). Prevalence and risk factors for anxiety symptoms during "
    "the outbreak of COVID-19: A large survey among 373216 junior and senior high "
    "school students in China. Journal of Affective Disorders, 288, 17–22. "
    "https://doi.org/10.1016/j.jad.2021.03.080")))
r.append(("TLTK Brunborg", whole(d1, "Brunborg, G. S., Skogen, J. C., Nilsen",
    "Brunborg, G. S., Nilsen, S. A., Skogen, J. C., & Bang, L. (2025). Possible "
    "explanations for the upward trend in mental distress among adolescents in Norway "
    "from 2011 to 2024. Social Science & Medicine, 384, Article 118528. "
    "https://doi.org/10.1016/j.socscimed.2025.118528")))
r.append(("body Lee→2026", part(d1, "Lee và cộng sự (2025)", "Lee và cộng sự (2026)")))
r.append(("body UNICEF paren", part(d1, "(UNICEF Việt Nam, 2022)",
                                    "(Viện Xã hội học và cộng sự, 2022)")))
r.append(("body UNICEF báo cáo", part(d1, "Báo cáo của UNICEF Việt Nam (2022)",
                                       "Báo cáo của Viện Xã hội học và cộng sự (2022)")))
d1.save(B1)
print("BÀI 1:")
for k, v in r:
    print(f"  {k}: {v}")
left = sum(p.text.count("UNICEF Việt Nam") for p in d1.paragraphs)
print(f"  còn 'UNICEF Việt Nam' trong body: {left}")

# ===== BÀI 2 =====
d2 = Document(B2)
r = []
r.append(("TLTK Li", whole(d2, "Li, S. H., Achilles, M. R., Werner-Seidler",
    "Li, S. H., Achilles, M. R., Spanos, S., Habak, S., Werner-Seidler, A., & O'Dea, B. "
    "(2022). A cognitive behavioural therapy smartphone app for adolescent depression "
    "and anxiety: Co-design of ClearlyMe. The Cognitive Behaviour Therapist, 15, "
    "Article e13. https://doi.org/10.1017/S1754470X22000095")))
r.append(("TLTK UNICEF", whole(d2, "UNICEF Vietnam. (2022)", UNICEF_NEW)))
r.append(("body UNICEF báo cáo", part(d2, "Báo cáo của UNICEF Việt Nam (2022)",
                                       "Báo cáo của Viện Xã hội học và cộng sự (2022)")))
# in-text Li ở khuyến nghị: "(Li và cộng sự, 2022)" — năm 2022 đúng, giữ nguyên
d2.save(B2)
print("BÀI 2:")
for k, v in r:
    print(f"  {k}: {v}")

# ===== BÀI 03 =====
d3 = Document(B03)
r = []
r.append(("TLTK Copeland", part(d3,
    "Longitudinal patterns of anxiety from childhood to adulthood. Journal",
    "Longitudinal patterns of anxiety from childhood to adulthood: The Great Smoky "
    "Mountains Study. Journal")))
r.append(("TLTK Essau", part(d3, "Moon-Ho, R. H.", "Ho, M.-H. R.")))
d3.save(B03)
print("BÀI 03:")
for k, v in r:
    print(f"  {k}: {v}")
print("[DONE]")
