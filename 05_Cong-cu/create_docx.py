"""
Script to create .docx files for each translated paper,
embedding PDF page images (tables, figures) from originals.
Also creates the summary table as .docx.
"""
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
import glob

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
TRANS_DIR = os.path.join(BASE, "Translations")
ORIG_DIR = os.path.join(BASE, "Originals")
OUT_DIR = os.path.join(BASE, "DocFiles")
IMG_DIR = os.path.join(BASE, "PDF_Images")

os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(IMG_DIR, exist_ok=True)


def extract_pdf_pages_as_images(pdf_path, output_prefix):
    """Extract each page of PDF as a PNG image."""
    images = []
    try:
        doc = fitz.open(pdf_path)
        for i, page in enumerate(doc):
            mat = fitz.Matrix(2.0, 2.0)  # 2x zoom for quality
            pix = page.get_pixmap(matrix=mat)
            img_path = os.path.join(IMG_DIR, f"{output_prefix}_page_{i+1}.png")
            pix.save(img_path)
            images.append(img_path)
        doc.close()
    except Exception as e:
        print(f"  Error extracting {pdf_path}: {e}")
    return images


def add_heading_styled(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    return h


def add_para(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.size = Pt(size)
    run.bold = bold
    return p


def md_to_docx(md_path, docx_path, pdf_images=None):
    """Convert a markdown translation file to docx, appending PDF page images."""
    doc = Document()

    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_critical_review = False  # Track if we're in the critical review section

    for line in lines:
        line = line.rstrip('\n')
        if not line.strip():
            continue

        # Detect critical review section
        if 'QUAN ĐIỂM PHẢN BIỆN' in line or 'CRITICAL REVIEW' in line:
            in_critical_review = True

        # Handle headings
        if line.startswith('# '):
            add_heading_styled(doc, line[2:].strip(), level=1)
        elif line.startswith('## '):
            h = add_heading_styled(doc, line[3:].strip(), level=2)
            if in_critical_review:
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('### '):
            add_heading_styled(doc, line[4:].strip(), level=3)
        elif line.startswith('#### '):
            add_heading_styled(doc, line[5:].strip(), level=4)
        elif line.startswith('---'):
            doc.add_paragraph('─' * 60)
        elif line.startswith('| ') and '|' in line[1:]:
            # Table row - add as paragraph with monospace feel
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(9)
            run.font.name = 'Consolas'
        elif line.startswith('- **') or line.startswith('  - '):
            # Bold list item
            p = doc.add_paragraph(style='List Bullet')
            # Parse bold markers
            text = line.lstrip('- ')
            if '**' in text:
                parts = text.split('**')
                for j, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                        if j % 2 == 1:  # odd indices are bold
                            run.bold = True
                        if in_critical_review:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                if in_critical_review:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            if in_critical_review:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('**') and line.endswith('**'):
            p_obj = add_para(doc, line.strip('*'), bold=True)
            if in_critical_review:
                for run in p_obj.runs:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        else:
            # Regular paragraph - handle inline bold
            p = doc.add_paragraph()
            text = line
            if '**' in text:
                parts = text.split('**')
                for j, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(12)
                        if j % 2 == 1:
                            run.bold = True
                        if in_critical_review:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(text)
                run.font.size = Pt(12)
                if in_critical_review:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Append PDF page images at the end
    if pdf_images:
        doc.add_page_break()
        add_heading_styled(doc, "BẢN GỐC / ORIGINAL PAPER PAGES", level=1)
        add_para(doc, "Các trang từ bản gốc PDF được đính kèm bên dưới để tham khảo và đối chiếu.", size=11)

        for img_path in pdf_images:
            try:
                doc.add_picture(img_path, width=Inches(6.5))
                last_paragraph = doc.paragraphs[-1]
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            except Exception as e:
                add_para(doc, f"[Lỗi chèn hình: {e}]")

    doc.save(docx_path)
    print(f"  Created: {docx_path}")


def create_summary_docx():
    """Create the summary table as a docx file."""
    md_path = os.path.join(BASE, "BANG_TOM_TAT_11_BAI_BAO.md")
    docx_path = os.path.join(OUT_DIR, "BANG_TOM_TAT_TONG_HOP.docx")

    if os.path.exists(md_path):
        md_to_docx(md_path, docx_path)
        print(f"  Summary created: {docx_path}")


# Map translation files to their PDF originals
pdf_map = {
    "01_Jenkins_et_al_2023": "01_Jenkins_et_al_2023.pdf",
    "02_Saikia_et_al_2023": "02_Saikia_et_al_2023.pdf",
    "03_Mandaknalli_Malusare_2021": "03_Mandaknalli_Malusare_2021.pdf",
    "04_NSCH_2020": None,  # national survey
    "05_Alharbi_et_al_2019": "05_Alharbi_et_al_2019.pdf",
    "06_Nakie_et_al_2022": "06_Nakie_et_al_2022.pdf",
    "07_Chen_et_al_2023": "07_Chen_et_al_2023.pdf",
    "08_Wen_et_al_2020": "08_Wen_et_al_2020.pdf",
    "09_Qiu_et_al_2022": "09_Qiu_et_al_2022.pdf",
    "10_Xu_et_al_2021": "10_Xu_et_al_2021.pdf",
    "11_Bhardwaj_et_al_2020": "11_Bhardwaj_et_al_2020.pdf",
}

print("=" * 60)
print("CREATING DOCX FILES WITH PDF IMAGES")
print("=" * 60)

# Step 1: Extract PDF images
print("\n--- Extracting PDF page images ---")
pdf_images_map = {}
for key, pdf_name in pdf_map.items():
    if pdf_name:
        pdf_path = os.path.join(ORIG_DIR, pdf_name)
        if os.path.exists(pdf_path):
            print(f"  Extracting: {pdf_name}")
            images = extract_pdf_pages_as_images(pdf_path, key)
            pdf_images_map[key] = images
            print(f"    -> {len(images)} pages extracted")

# Step 2: Create DOCX for each translation
print("\n--- Creating DOCX translations ---")
for md_file in sorted(glob.glob(os.path.join(TRANS_DIR, "*.md"))):
    basename = os.path.splitext(os.path.basename(md_file))[0]
    docx_path = os.path.join(OUT_DIR, f"{basename}.docx")
    images = pdf_images_map.get(basename, None)
    print(f"  Processing: {basename}")
    md_to_docx(md_file, docx_path, images)

# Step 3: Create summary DOCX
print("\n--- Creating Summary DOCX ---")
create_summary_docx()

print("\n" + "=" * 60)
print("DONE! All files saved to:", OUT_DIR)
print("=" * 60)
