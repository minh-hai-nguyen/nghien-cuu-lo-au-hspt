# -*- coding: utf-8 -*-
"""Convert jp2 image to png, then insert into docx after Trang 6 marker."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from PIL import Image

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/unicef_imgs/'

src = os.path.join(IMG_DIR, 'p006_img1_Im0.jp2')
dst = os.path.join(IMG_DIR, 'p006_img1_Im0.png')

try:
    im = Image.open(src)
    # Convert to RGB if needed
    if im.mode not in ('RGB', 'L'):
        im = im.convert('RGB')
    im.save(dst, 'PNG')
    print(f'Converted: {dst} ({os.path.getsize(dst):,} bytes)')
except Exception as e:
    print(f'Conversion failed: {e}')
    sys.exit(1)

d = Document(PATH)

# Find Trang 6 marker
anchor = None
for p in d.paragraphs:
    if re.match(r'---\s*Trang\s+6\b', p.text.strip()):
        anchor = p
        break

if anchor is None:
    print('Trang 6 marker not found!')
    sys.exit(1)

# Add image at end, then move
img_p = d.add_paragraph()
img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
img_p.add_run().add_picture(dst, width=Cm(11.0))

cap_p = d.add_paragraph()
cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = cap_p.add_run('Hình 2: Ảnh minh hoạ đầu chương Lời nói đầu\n(Figure 2: Foreword chapter opener illustration — p.6)')
r.font.name = 'Times New Roman'
r.font.size = Pt(9)
r.italic = True
r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

# Move both elements to right after anchor
img_el = img_p._element
cap_el = cap_p._element
anchor_el = anchor._element

img_el.getparent().remove(img_el)
cap_el.getparent().remove(cap_el)

anchor_el.addnext(cap_el)
anchor_el.addnext(img_el)

d.save(PATH)
print('Trang 6 image inserted.')

import zipfile
with zipfile.ZipFile(PATH) as z:
    media = [n for n in z.namelist() if n.startswith('word/media/')]
    print(f'Images now in docx: {len(media)}')
    for m in sorted(media):
        info = z.getinfo(m)
        print(f'  {m} ({info.file_size:,} bytes)')
