"""Doc trả lời: t-statistic là gì + tác dụng - format mới (câu trả lời tô xanh trước phụ lục)."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/t_statistic_dien_giai_va_tac_dung.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
BLUE = RGBColor(0, 112, 192)
DARK = RGBColor(31, 73, 125)
GREEN = RGBColor(54, 95, 44)
RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(191, 97, 14)
GRAY = RGBColor(90, 90, 90)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def vn_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p
def blue_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = BLUE
    return p

# ===== TITLE =====
title = d.add_heading('Hệ số t (t-statistic): Là gì, dùng để làm gì?', level=0)
for r in title.runs: r.font.color.rgb = DARK
sub = d.add_paragraph()
sr = sub.add_run('Diễn giải t-statistic — chỉ số kiểm định phổ biến nhất trong nghiên cứu lo âu')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = GRAY
d.add_paragraph()

# ===== CÂU HỎI =====
qbox = d.add_table(rows=1, cols=1); qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]; shade(cell, 'FFF8DC')
p = cell.paragraphs[0]
r = p.add_run('Câu hỏi của thầy:'); r.bold = True
p2 = cell.add_paragraph()
p2.add_run('"Có một số dữ liệu đưa ra hệ số t. Thầy không biết t có tác dụng gì, Em ơi?"')
d.add_paragraph()

# ===== BỐI CẢNH =====
add_h('Bối cảnh', 1)
vn_para('Khi đọc các bài nghiên cứu lo âu, thầy thường gặp các con số như: t = 2,6 (p = 0,01); t₃₆ = 2,6; '
        't(124) = 3,45 (p < 0,001). Những "t" này gọi là HỆ SỐ T (t-statistic) — chỉ số quan trọng của KIỂM ĐỊNH T '
        '(t-test) — một trong những công cụ thống kê PHỔ BIẾN NHẤT trong nghiên cứu tâm lý + dịch tễ.')
vn_para('Trong CSDL của em, t-statistic xuất hiện ở: QT062 Clear Fear (t₃₆=2,6, p=0,01); QT001 Jenkins (Mann-Whitney '
        'thay t-test do mẫu nhỏ); EACP QT065 (t-test cho post-test); và nhiều bài khác.')
d.add_paragraph()

# ===== KIẾN THỨC NỀN =====
add_h('Kiến thức nền', 1)

vn_para('1. t-statistic là gì?', bold=True, size=12)
vn_para('t = (giá trị quan sát − giá trị giả thuyết H₀) / sai số chuẩn (SE) của giá trị quan sát.')
vn_para('Nói cách khác: t đo "khoảng cách" giữa kết quả thực tế và giả thuyết H₀ (giả thuyết "không có khác biệt" hoặc '
        '"không có hiệu ứng"), được CHUẨN HOÁ theo độ biến thiên của dữ liệu (SE).')
d.add_paragraph()

vn_para('2. Ý nghĩa trực quan của t', bold=True, size=12)
intuition_tbl = d.add_table(rows=4, cols=2); intuition_tbl.style = 'Table Grid'
intuition_data = [
    ('|t| = 0', 'Kết quả CHÍNH XÁC bằng giả thuyết H₀ → không có khác biệt'),
    ('|t| ≈ 1', 'Khác biệt NHỎ — có thể chỉ do ngẫu nhiên'),
    ('|t| ≈ 2', 'Khác biệt VỪA — bắt đầu nghi ngờ H₀ (thường p ≈ 0,05)'),
    ('|t| ≥ 3', 'Khác biệt LỚN — gần như chắc chắn không phải ngẫu nhiên (p < 0,01)'),
]
for i, (k, v) in enumerate(intuition_data):
    c0 = intuition_tbl.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p0 = c0.paragraphs[0]; r0 = p0.add_run(k); r0.bold = True
    intuition_tbl.rows[i].cells[1].text = v
d.add_paragraph()

vn_para('3. Tác dụng của t-statistic — 3 mục đích chính', bold=True, size=12)

purpose_tbl = d.add_table(rows=4, cols=2); purpose_tbl.style = 'Table Grid'
hdr = ['Mục đích', 'Giải thích + Ví dụ']
for i, h in enumerate(hdr):
    c = purpose_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
purposes = [
    ('① So sánh 2 trung bình (t-test 2 mẫu)',
     'Vd: so sánh điểm GAD-7 nam vs nữ. Nếu t = 2,8 (p < 0,01): có khác biệt thực giữa 2 nhóm. '
     'CHỈ DÙNG khi: (a) data phân phối chuẩn; (b) phương sai 2 nhóm tương đồng (homogeneity); '
     '(c) cỡ mẫu vừa đủ. Nếu không thoả → dùng Mann-Whitney U thay'),
    ('② So sánh trước-sau (paired t-test)',
     'Vd: điểm GAD-7 cùng 1 HS trước và sau can thiệp 8 tuần. Nếu t = 3,5 (p < 0,001): can thiệp có hiệu lực. '
     'Đây là test sử dụng nhiều nhất trong RCT đơn nhóm + pre-post design (vd QT062 Clear Fear: t₃₆=2,6, p=0,01)'),
    ('③ So sánh 1 trung bình với giá trị chuẩn (one-sample t-test)',
     'Vd: điểm GAD-7 trung bình mẫu HS VN có khác với ngưỡng 5 (mức nhẹ) không? '
     'Nếu t = 4,2 (p < 0,001): có khác. Ít dùng hơn 2 cách trên'),
]
for i, (k, v) in enumerate(purposes):
    c0 = purpose_tbl.rows[i+1].cells[0]; shade(c0, 'E2EFDA')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = GREEN
    purpose_tbl.rows[i+1].cells[1].text = v
d.add_paragraph()

vn_para('4. Cách viết t trong báo cáo', bold=True, size=12)
vn_para('Format chuẩn APA 7th: t(df) = giá_trị, p = p-value', bold=True)
format_tbl = d.add_table(rows=4, cols=2); format_tbl.style = 'Table Grid'
fdata = [
    ('t(36) = 2,6, p = 0,01', '36 = degrees of freedom (df) — thường là n − 1 hoặc n − 2 tuỳ test'),
    ('t₃₆ = 2,6, p = 0,01', 'df viết dưới dạng subscript — kiểu khác của t(36)'),
    ('t = 2,6, p = 0,01', 'Một số tạp chí bỏ df — cách viết tắt'),
    ('t(36) = 2,6, p < 0,001', 'p < 0,001: rất có ý nghĩa thống kê (so với p = 0,01)'),
]
for i, (k, v) in enumerate(fdata):
    c0 = format_tbl.rows[i].cells[0]; shade(c0, 'D9E1F2')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True
    format_tbl.rows[i].cells[1].text = v
d.add_paragraph()

vn_para('5. df (degrees of freedom — bậc tự do)', bold=True, size=12)
vn_para('df ảnh hưởng cách giải thích t — cùng giá trị t, nhưng df khác sẽ cho p khác:')
df_tbl = d.add_table(rows=4, cols=3); df_tbl.style = 'Table Grid'
hdr = ['t = 2,0', 'df', 'p (2-tailed)']
for i, h in enumerate(hdr):
    c = df_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
ddata = [
    ('', 'df = 5 (mẫu rất nhỏ)', 'p ≈ 0,10 — KHÔNG có ý nghĩa'),
    ('', 'df = 30', 'p ≈ 0,054 — sát ngưỡng'),
    ('', 'df = 200 (mẫu lớn)', 'p ≈ 0,047 — CÓ ý nghĩa'),
]
for i, (a, b, c) in enumerate(ddata):
    df_tbl.rows[i+1].cells[0].text = a
    df_tbl.rows[i+1].cells[1].text = b
    df_tbl.rows[i+1].cells[2].text = c
    for pp in df_tbl.rows[i+1].cells[0].paragraphs:
        for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('→ Quy tắc: cùng t = 2,0, mẫu CÀNG LỚN (df cao) thì p càng nhỏ — kết quả "có ý nghĩa" hơn. '
        'Đó là vì với mẫu lớn, ngẫu nhiên ít có khả năng tạo ra chênh lệch lớn.', bold=True)
d.add_paragraph()

vn_para('6. t vs F vs Mann-Whitney — khi nào dùng cái nào?', bold=True, size=12)
which_tbl = d.add_table(rows=5, cols=3); which_tbl.style = 'Table Grid'
hdr = ['Test', 'Khi nào dùng', 'Ví dụ']
for i, h in enumerate(hdr):
    c = which_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
wdata = [
    ('t-test', '2 nhóm, data phân phối chuẩn, phương sai tương đồng',
     'GAD-7 nam vs nữ (n>30 mỗi nhóm)'),
    ('Mann-Whitney U', '2 nhóm, data KHÔNG chuẩn hoặc mẫu nhỏ',
     'Jenkins QT001 (n=75): GAD-10 nam vs nữ (p=0,016)'),
    ('ANOVA / F-test', '≥ 3 nhóm so sánh, data chuẩn',
     'Bảng RLLATQ × 4 khối lớp (lớp 6, 7, 8, 9): F = 44,484'),
    ('Paired t-test', 'Cùng 1 người, 2 thời điểm (trước-sau)',
     'Clear Fear QT062: t₃₆=2,6, p=0,01 (GAD-7 baseline vs follow-up)'),
]
for i, (k, w, e) in enumerate(wdata):
    c0 = which_tbl.rows[i+1].cells[0]; shade(c0, 'FFE699')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = ORANGE
    which_tbl.rows[i+1].cells[1].text = w
    which_tbl.rows[i+1].cells[2].text = e
d.add_paragraph()

vn_para('7. t vs effect size — KHÔNG đồng nhất', bold=True, size=12)
warn_tbl = d.add_table(rows=1, cols=1); warn_tbl.style = 'Table Grid'
wc = warn_tbl.rows[0].cells[0]; shade(wc, 'FCE4D6')
wp = wc.paragraphs[0]
wr = wp.add_run('⚠ LỖI PHỔ BIẾN: ')
wr.bold = True; wr.font.color.rgb = RED
wp.add_run('"t lớn = effect size lớn" — KHÔNG ĐÚNG. t phụ thuộc CẢ hiệu ứng VÀ cỡ mẫu. '
           'Mẫu CỰC LỚN có thể cho t lớn cho cả khác biệt nhỏ. Phải đọc t kèm Cohen\'s d (effect size). '
           'Vd: Clear Fear t=2,6 trong khi Cohen\'s d ≈ 0,4-0,5 (medium effect). Một bài khác có t=10,0 chưa chắc d > 1 nếu n=10.000.')
d.add_paragraph()

vn_para('Cách convert t sang Cohen\'s d (gần đúng):', bold=True)
vn_para('• Paired t-test: d ≈ t / √n')
vn_para('• Independent t-test 2 nhóm: d ≈ 2t / √(df) (gần đúng)')
vn_para('Ví dụ Clear Fear: t₃₆ = 2,6 → d ≈ 2,6 / √36 = 2,6 / 6 = 0,43 → MEDIUM effect')
d.add_paragraph()

# =====================================================
# CÂU TRẢ LỜI (TÔ XANH, GOM TRƯỚC PHỤ LỤC)
# =====================================================
add_h('Câu trả lời', 1, BLUE)

blue_para('Tác dụng của hệ số t — tóm gọn:', bold=True, size=12)
blue_para('t-statistic là CHỈ SỐ KIỂM ĐỊNH dùng để: (1) so sánh 2 nhóm; (2) so sánh trước-sau; (3) so sánh 1 mẫu với giá trị chuẩn. '
          't = (chênh lệch quan sát) / (sai số chuẩn). |t| càng lớn → khác biệt càng rõ → ít có khả năng do ngẫu nhiên.')
d.add_paragraph()

blue_para('Quy tắc đọc nhanh:', bold=True, size=12)
read_tbl = d.add_table(rows=5, cols=2); read_tbl.style = 'Table Grid'
read_hdr = ['|t|', 'Diễn giải']
for i, h in enumerate(read_hdr):
    c = read_tbl.rows[0].cells[i]; shade(c, 'BDD7EE')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = BLUE
rdata = [
    ('|t| ≈ 0', 'KHÔNG khác biệt'),
    ('|t| ≈ 1', 'Khác biệt nhỏ — có thể ngẫu nhiên'),
    ('|t| ≈ 2', 'Khác biệt vừa — sát ngưỡng có ý nghĩa (p ≈ 0,05)'),
    ('|t| ≥ 3', 'Khác biệt RÕ — gần chắc chắn không ngẫu nhiên (p < 0,01)'),
]
for i, (k, v) in enumerate(rdata):
    c = read_tbl.rows[i+1].cells[0]
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(k); rr.font.color.rgb = BLUE; rr.bold = True
    c2 = read_tbl.rows[i+1].cells[1]
    pp2 = c2.paragraphs[0]; rr2 = pp2.add_run(v); rr2.font.color.rgb = BLUE
d.add_paragraph()

blue_para('3 LOẠI t-test — khi nào dùng:', bold=True, size=12)
for it in [
    'Independent t-test (2 nhóm độc lập): so điểm GAD-7 nam vs nữ',
    'Paired t-test (cùng người, 2 thời điểm): so điểm GAD-7 trước vs sau can thiệp',
    'One-sample t-test (1 mẫu vs giá trị chuẩn): so điểm GAD-7 mẫu với ngưỡng 5',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Đọc t kèm 3 thông tin khác:', bold=True, size=12)
for it in [
    'p-value: Có ý nghĩa thống kê hay không (p < 0,05)',
    'df (degrees of freedom): Cùng t, df khác → p khác. df = n−1 (paired) hoặc n−2 (independent)',
    'Effect size (Cohen\'s d): t lớn KHÔNG có nghĩa effect size lớn. Mẫu lớn dễ làm t lớn dù khác biệt nhỏ',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Cảnh báo quan trọng:', bold=True, size=12)
blue_para('"t lớn = effect size lớn" KHÔNG ĐÚNG. t phụ thuộc CẢ effect size VÀ cỡ mẫu. Phải đọc t cùng Cohen\'s d. '
          'Convert nhanh: paired d ≈ t/√n; independent d ≈ 2t/√df.')
d.add_paragraph()

blue_para('Khi nào KHÔNG dùng t-test:', bold=True, size=12)
for it in [
    'Mẫu nhỏ + data không chuẩn → dùng Mann-Whitney U thay (vd Jenkins QT001 n=75)',
    '≥ 3 nhóm so sánh → dùng ANOVA (F-test) thay (vd RLLA × 4 khối lớp)',
    'Outcome nhị phân (có/không) → dùng chi-square hoặc logistic regression thay',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Ví dụ thực tế từ DB của em:', bold=True, size=12)
blue_para('Clear Fear app QT062 (Samele 2025 UK): t₃₆ = 2,6; p = 0,01. Diễn giải: GAD-7 sau 9 tuần dùng app GIẢM '
          'có ý nghĩa thống kê so với baseline (paired t-test, df=36 vì n=37). Effect size ước tính d ≈ 0,43 — '
          'medium effect. Không phải hiệu ứng lớn nhưng có ý nghĩa lâm sàng.')
d.add_paragraph()

blue_para('Ghi nhớ 1 dòng:', bold=True, size=12)
blue_para('t-statistic = chỉ số đo "chênh lệch chuẩn hoá". |t| ≥ 2 thường có ý nghĩa thống kê. '
          'Đọc t kèm p, df, và Cohen\'s d. KHÔNG đánh đồng "t lớn" với "effect size lớn".', bold=True)

d.add_paragraph()

# =====================================================
# PHỤ LỤC
# =====================================================
add_h('Phụ lục — Tài liệu tham khảo', 1)
refs = [
    'Student (William Sealy Gosset, 1908). The probable error of a mean. Biometrika, 6(1):1-25. [Bài gốc giới thiệu t-distribution]',
    'Cohen J. (1988). Statistical Power Analysis for the Behavioral Sciences (2nd ed.). Routledge. [Cohen\'s d: 0,2 small / 0,5 medium / 0,8 large]',
    'Field A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). Sage. [Giáo khoa t-test]',
    'Lakens D. (2013). Calculating and reporting effect sizes to facilitate cumulative science. Frontiers in Psychology, 4:863. [Convert t → d]',
    'Samele C, Urquia N, Edwards R, Donnell K, Krause N. (2025). Evaluation of the Clear Fear Smartphone App. JMIR Formative Research. DOI 10.2196/55603. [QT062 — ví dụ t₃₆=2,6, p=0,01]',
    'Jenkins JH, et al. (2023). Depression and anxiety among multiethnic middle school students. [QT001 — ví dụ Mann-Whitney thay t-test do mẫu nhỏ]',
    'Doc liên quan trong dự án: 01_Bao-cao/F_statistic_ANOVA_dien_giai.docx (so sánh ≥3 nhóm); 01_Bao-cao/NNT_va_MannWhitney_chenh_lech_trung_vi.docx (Mann-Whitney khi data không chuẩn).',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs: r.font.size = Pt(10)

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
