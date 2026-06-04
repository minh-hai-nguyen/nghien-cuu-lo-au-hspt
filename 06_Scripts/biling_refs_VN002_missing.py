# -*- coding: utf-8 -*-
"""Add VN notes for Vietnamese-origin references (5 missing)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor

OUT = '03_Ban-dich/VN002_VNAMHS_2022_National_FULL.docx'

# Pattern prefix → EN note (Vietnamese-origin, no translation needed)
VN_NOTE_MAP = {
    'Bộ Giáo dục và Đào tạo. 2022a': '(Văn bản tiếng Việt gốc — không cần dịch. Đã bản địa hoá trước khi báo cáo xuất bản.)',
    'Bộ Giáo dục và Đào tạo. 2022b': '(Văn bản tiếng Việt gốc — không cần dịch.)',
    'Bộ Y tế. 2022': '(Văn bản tiếng Việt gốc — không cần dịch.)',
    'Cao Tiến Đức. 2020': '(Tác phẩm tiếng Việt gốc — không cần dịch.)',
    'Chính phủ. 2022': '(Văn bản tiếng Việt gốc — không cần dịch.)',
    'Hoàng Thế Hải and Bùi Thị Thanh Diệu. 2021': '(Tác phẩm tiếng Việt gốc — không cần dịch.)',
    'MOLISA (Bộ Lao động-Thương binh và Xã hội). 2012': '(Văn bản tiếng Việt gốc — không cần dịch.)',
    'Thủ tướng Chính phủ. 2011': '(Văn bản tiếng Việt gốc — không cần dịch.)',
    'Thủ tướng Chính phủ. 2020': '(Văn bản tiếng Việt gốc — không cần dịch.)',
    'Thủ tướng Chính phủ. 2022': '(Văn bản tiếng Việt gốc — không cần dịch.)',
}

d = Document(OUT)

added = 0
for prefix, note in VN_NOTE_MAP.items():
    for i, p in enumerate(d.paragraphs):
        if p.text.startswith(prefix):
            # Check if next paragraph already has the note
            nxt = d.paragraphs[i+1].text if i+1 < len(d.paragraphs) else ''
            if 'không cần dịch' in nxt:
                continue  # already added
            # Add new paragraph after this one
            new_para = d.add_paragraph()
            r = new_para.add_run(note)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            # Move it right after current paragraph
            en_el = p._element
            new_el = new_para._element
            new_el.getparent().remove(new_el)
            en_el.addnext(new_el)
            print(f'  + {prefix[:40]:45s} added note')
            added += 1
            break

print(f'\nTotal VN-origin notes added: {added}')
d.save(OUT)
