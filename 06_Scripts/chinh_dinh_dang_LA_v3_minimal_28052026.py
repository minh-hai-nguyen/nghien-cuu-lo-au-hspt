# -*- coding: utf-8 -*-
"""Chinh dinh dang LA v3 — MINIMAL CHANGES:
- CHI lam: margins (section), TOC field, xoa watermark + metadata
- KHONG dong cham bat ky paragraph hay run nao
- Giu nguyen 100% noi dung + dinh dang van ban
28/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_6_FixAnderson_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

print(f"SRC: {SRC}")
print(f"OUT: {OUT}")
print()
d = Document(SRC)

# ============================================================
# 1. Margins
# ============================================================
print("--- Margins ---")
for sec in d.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2.0)
print(f"  Set 2.5/2.5/3.5/2.0 on {len(d.sections)} sections")
print()

# ============================================================
# 2. Watermark + headers/footers
# ============================================================
print("--- Watermarks ---")
removed = 0
for sec in d.sections:
    for hf in [sec.header, sec.first_page_header, sec.even_page_header,
               sec.footer, sec.first_page_footer, sec.even_page_footer]:
        try:
            for elem in list(hf._element.iter()):
                if elem.tag in (qn('w:pict'), qn('w:object'), qn('w:drawing')):
                    if elem.getparent() is not None:
                        elem.getparent().remove(elem)
                        removed += 1
        except: pass
print(f"  Removed {removed} watermark elements from H/F")
print()

# ============================================================
# 3. TOC field — chen sau "MỤC LỤC" co san
# ============================================================
print("--- TOC field ---")
def make_toc_paragraph(parent_para):
    r = parent_para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r._r.append(instrText)
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fldChar_sep)
    instrResult = OxmlElement('w:t')
    instrResult.text = 'Nhấn F9 hoặc chuột phải > Update Field để cập nhật Mục lục.'
    r._r.append(instrResult)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar_end)

toc_idx = None
for i, p in enumerate(d.paragraphs[:200]):
    if p.text.strip().upper() in ('MỤC LỤC', 'MUC LUC'):
        toc_idx = i
        break

if toc_idx is not None:
    toc_heading_p = d.paragraphs[toc_idx]
    new_p_xml = OxmlElement('w:p')
    toc_heading_p._element.addnext(new_p_xml)
    from docx.text.paragraph import Paragraph
    toc_p_obj = Paragraph(new_p_xml, toc_heading_p._parent)
    make_toc_paragraph(toc_p_obj)
    print(f"  TOC inserted after para {toc_idx} (MỤC LỤC)")
else:
    print(f"  No MỤC LỤC found - TOC NOT inserted")
print()

# ============================================================
# 4. Metadata
# ============================================================
print("--- Metadata ---")
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
print("  Cleared")
print()

d.save(OUT)
print(f"--- Saved: {OUT} ({os.path.getsize(OUT)//1024} KB) ---")
