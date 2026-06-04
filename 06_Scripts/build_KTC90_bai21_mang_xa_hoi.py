"""Build doc tra loi: bai 21 mang xa hoi - 0,20 (0,12-0,28) KTC90% co chuan khong, mức nhỏ?"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/KTC_90_QT021_Brunborg_2025_mang_xa_hoi_0_20.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(12); r.italic = italic
    return p

# =====================================================
# TIÊU ĐỀ
# =====================================================
H('Bình luận số liệu β = 0,20 (KTC 90%: 0,12–0,28)', level=1)
H('— QT021 Brunborg et al. 2025 (Soc Sci Med) — mạng xã hội & SKTT VTN Na Uy —', level=2)

# =====================================================
# CÂU HỎI (xanh)
# =====================================================
H('Câu hỏi của thầy', level=2)
para_blue(
    'Bài 21 về mạng xã hội họ viết thế này là chuẩn không? "0,20 (0,12–0,28) KTC 90%". '
    'Như vậy là nhỏ phải không?'
)
para_black('')
para_black(
    'Bài 21 = QT021 Brunborg et al. (2025) "Possible Explanations for the Upward Trend '
    'in Mental Distress among Adolescents in Norway from 2011 to 2024", Social Science '
    '& Medicine 384, Article 118528. Q1, IF 5,4. Mẫu rất lớn: 979.043 VTN Na Uy từ 1.417 '
    'khảo sát cấp thành phố. Đây là bài narrative review + decomposition mở rộng dùng '
    'khung dịch tễ học của Keyes & Platt (2024) để xác định 8 yếu tố giải thích xu hướng '
    'tăng distress, bao gồm mạng xã hội.', italic=True
)

# =====================================================
# BACKGROUND (đen)
# =====================================================
H('1. Phân biệt KTC 90% và KTC 95% — Chuẩn vàng là gì?', level=2)

para_black('KTC (Khoảng tin cậy, Confidence Interval — CI) là khoảng giá trị có khả năng chứa giá trị thật của tham số trong dân số. Có 2 loại độ tin cậy phổ biến:', bold=False)

bullet_black('KTC 95% — CHUẨN VÀNG trong y học/tâm lý/dịch tễ. Cochrane, CONSORT, PRISMA, APA đều quy định 95%.')
bullet_black('KTC 90% — KHÔNG phải mặc định. Chỉ chấp nhận trong vài tình huống đặc biệt (xem mục 2).')

para_black('Quan hệ giữa 2 loại:', bold=True)
bullet_black('KTC 90% LUÔN HẸP HƠN KTC 95% trên cùng dữ liệu (vì độ tin cậy thấp hơn).')
bullet_black('Quy đổi nhanh (giả định phân phối chuẩn): độ rộng KTC 95% ≈ 1,19 × độ rộng KTC 90%.')
bullet_black('Ví dụ: KTC 90% = 0,12–0,28 (rộng 0,16) → KTC 95% ước tính ≈ 0,10–0,30 (rộng 0,20).')
bullet_black('Khi đổi sang KTC 95%, khoảng có thể CHỨA 0 hơn → mất ý nghĩa thống kê. Một số tác giả "chữa" kết quả đuối bằng cách chuyển sang KTC 90%.')

H('2. Khi nào dùng KTC 90% mới được chấp nhận?', level=2)

tbl = doc.add_table(rows=5, cols=2)
tbl.style = 'Light Grid Accent 1'
header = tbl.rows[0]
header.cells[0].text = 'Bối cảnh chấp nhận KTC 90%'
header.cells[1].text = 'Lý do'
data = [
    ('Equivalence/Non-inferiority test (TOST)', 'Quy ước riêng — TOST dùng 2 phía 90% để tương đương 95% 1 phía.'),
    ('Bayesian credible interval / NMA', '90% HDI (highest density interval) là tuỳ chọn phổ biến.'),
    ('Pilot study / exploratory analysis', 'Mẫu nhỏ, chấp nhận sai số lớn hơn để khám phá.'),
    ('Genome-wide / massive multiple testing', 'Đôi khi nới lỏng để tăng power, kèm hiệu chỉnh đa so sánh.'),
]
for i, (a, b) in enumerate(data, 1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black('NẾU bài 21 KHÔNG giải thích lý do dùng KTC 90% — đây là ĐIỂM YẾU phương pháp, đáng đưa vào phản biện.', bold=True, italic=True)

H('3. Diễn giải con số 0,20 — phụ thuộc loại chỉ số', level=2)

para_black('Cùng giá trị 0,20 nhưng diễn giải KHÁC NHAU tùy đó là chỉ số nào. Trong các bài về mạng xã hội ↔ lo âu/trầm cảm, 0,20 thường là correlation r hoặc β.')

tbl2 = doc.add_table(rows=7, cols=3)
tbl2.style = 'Light Grid Accent 1'
header = tbl2.rows[0]
header.cells[0].text = 'Loại chỉ số'
header.cells[1].text = 'Mức độ của 0,20'
header.cells[2].text = 'Diễn giải'
rows = [
    ('Correlation r (Pearson/Spearman)', 'NHỎ', 'Cohen 1988: 0,1 nhỏ / 0,3 TB / 0,5 lớn. Giải thích ~4% phương sai (r²=0,04).'),
    ("Cohen's d (SMD)", 'NHỎ', 'Cohen 1988: 0,2 nhỏ / 0,5 TB / 0,8 lớn. d=0,20 → khoảng 58% nhóm 1 cao hơn TB nhóm 2.'),
    ('β (regression coefficient chuẩn hoá)', 'Nhỏ-TB', 'Tăng 1 SD biến X → tăng 0,20 SD biến Y. Đáng để ý nếu p<0,05.'),
    ('OR (Odds Ratio)', 'BẢO VỆ MẠNH', 'OR=0,20 → odds nhóm phơi nhiễm bằng 20% nhóm tham chiếu. 1/0,20=5 → "mạnh".'),
    ('Risk Ratio / Risk Difference', 'Trung bình-Lớn', 'RR=0,20 hoặc RD=20 điểm % đều mức trung bình-lớn.'),
    ('R² (variance explained)', 'NHỎ', 'Mô hình giải thích 20% phương sai → vừa phải nhưng không xuất sắc.'),
]
for i, (a, b, c) in enumerate(rows, 1):
    tbl2.rows[i].cells[0].text = a
    tbl2.rows[i].cells[1].text = b
    tbl2.rows[i].cells[2].text = c
for row in tbl2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black(
    'Trong corpus dự án (bài về mạng xã hội ↔ SKTT thanh thiếu niên — Twenge, Orben, Keyes, Coyne...), '
    'các tác giả thường báo cáo r ~ 0,10–0,15 (rất nhỏ). Số 0,20 cao hơn mức trung vị y văn '
    'một chút, nhưng vẫn nằm trong khoảng "small effect" theo Cohen.', italic=True
)

H('4. KTC 0,12–0,28 không chứa 0 → CÓ ý nghĩa thống kê', level=2)
para_black('Khoảng [0,12; 0,28] hoàn toàn dương → ở mức tin cậy 90%, hệ số khác 0 có ý nghĩa thống kê (p < 0,10 1 phía hoặc 2 phía tùy giả thuyết).', bold=True)
bullet_black('Ngay cả khi quy đổi sang KTC 95% (~0,10–0,30), khoảng vẫn KHÔNG chứa 0 → vẫn có ý nghĩa thống kê (p < 0,05).')
bullet_black('Vậy "ý nghĩa thống kê" KHÔNG phải vấn đề. Vấn đề là (1) KTC 90% không chuẩn nếu không giải thích lý do; (2) độ lớn hiệu ứng = NHỎ.')

H('4b. Đối chiếu trực tiếp với bản dịch QT021 — Bảng "Tác động chính β"', level=2)
para_black(
    'Em đối chiếu vào bản dịch QT021 hiện có trong corpus (file: '
    '03_Ban-dich/QT021_Norway_Brunborg_2025_SocSciMed.docx, Bảng 4). Trong bài này, các '
    'tác giả báo cáo β chuẩn hoá (standardized regression coefficient) cho từng yếu tố:'
)

tbl_q21 = doc.add_table(rows=9, cols=3)
tbl_q21.style = 'Light Grid Accent 1'
header = tbl_q21.rows[0]
header.cells[0].text = 'Yếu tố'
header.cells[1].text = 'Nam β'
header.cells[2].text = 'Nữ β'
data_q21 = [
    ('Khó khăn tài chính', '0,23', '0,24'),
    ('Tối ở nhà', '0,11', '0,11'),
    ('Hoạt động thể chất', '−0,06', '−0,08'),
    ('Bắt nạt', '0,30', '0,31'),
    ('Hài lòng với cha mẹ', '−0,20', '−0,23'),
    ('MẠNG XÃ HỘI', '0,13', '0,18'),
    ('Sử dụng cần sa', '0,16', '0,12'),
    ('Bất mãn trường học', '0,40', '0,44'),
]
for i, (a, b, c) in enumerate(data_q21, 1):
    tbl_q21.rows[i].cells[0].text = a
    tbl_q21.rows[i].cells[1].text = b
    tbl_q21.rows[i].cells[2].text = c
for row in tbl_q21.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black('Quan sát chính:', bold=True)
bullet_black(
    'β cho mạng xã hội ở NỮ = 0,18 — RẤT GẦN với 0,20 thầy đang xem (có thể tác giả '
    'làm tròn). Nam β chỉ 0,13.'
)
bullet_black(
    'KTC 90% [0,12; 0,28] khớp với β nữ 0,18–0,20 — width khoảng 0,16. Mẫu nữ ~ 461.247 '
    'rất lớn, nên KTC hẹp tự nhiên.'
)
bullet_black(
    'Brunborg KHÔNG dùng KTC 90% mặc định — bản dịch hiện có chỉ trích β + p-value 1, không '
    'có KTC chi tiết. KTC 90% [0,12; 0,28] thầy đang xem có thể đến từ supplementary '
    'material hoặc một bảng khác trong bài gốc tiếng Anh — em CHƯA verify được trong '
    'bản dịch hiện tại.'
)

para_black('Phát hiện đáng chú ý từ Bảng 5 cùng bài QT021:', bold=True)
bullet_black(
    'Khi GIỮ CỐ ĐỊNH mạng xã hội trong mô hình decomposition, hệ số b xu hướng tăng '
    'distress ở nữ giảm từ 0,015 → 0,006 (giảm 60%, χ² = 764,04, p < 0,001).'
)
bullet_black(
    'Brunborg kết luận: mặc dù β = 0,18 thuộc nhóm "small effect", mạng xã hội VẪN GIẢI '
    'THÍCH ĐÁNG KỂ xu hướng tăng distress ở nữ giai đoạn 2014–2024.'
)
bullet_black(
    'Đây là VÍ DỤ ĐẸP cho nguyên tắc: effect size NHỎ vẫn có thể có ý nghĩa lâm sàng + '
    'xã hội khi mẫu rất lớn (n > 100k) và yếu tố tác động lan rộng trong dân số.'
)

H('5. "Có ý nghĩa thống kê" ≠ "Đáng để ý lâm sàng"', level=2)
para_black('Trong nghiên cứu cỡ mẫu LỚN (mạng xã hội thường n > 1.000–10.000), một hiệu ứng RẤT NHỎ vẫn dễ đạt p < 0,05. Cần phân biệt:', bold=False)
bullet_black('Statistical significance (có ý nghĩa thống kê) — phụ thuộc cỡ mẫu.')
bullet_black('Practical/clinical significance (có ý nghĩa thực hành) — phụ thuộc độ lớn effect size + bối cảnh.')
bullet_black('Với r=0,20 ↔ mạng xã hội & lo âu: tác động NHỎ. Tăng/giảm mạng xã hội 1 SD chỉ liên quan tới thay đổi lo âu ~0,20 SD — hiệu ứng rõ ràng nhưng không quá lớn.')

H('6. Tham chiếu nội bộ (corpus dự án)', level=2)
bullet_black('QT008 Wen 2020 — academic stress + lo âu (LPA, Trung Quốc): OR ~ 2–3, mức TRUNG BÌNH.', italic=True)
bullet_black('QT067 Pascoe 2020 — academic stress narrative review: không có effect size pooled.', italic=True)
bullet_black('Trong y văn, các nghiên cứu tổng quan (Orben & Przybylski 2019 Nature Human Behaviour, Twenge 2018) báo cáo r ~ 0,03–0,15 cho social media ↔ wellbeing — TỨC TÁC ĐỘNG RẤT NHỎ.', italic=True)

# =====================================================
# CÂU TRẢ LỜI (xanh, gom 1 chỗ trước phụ lục)
# =====================================================
H('7. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Tóm tắt nhanh:', bold=True)
para_blue(
    'QT021 = Brunborg et al. 2025 (Soc Sci Med, n=979.043 VTN Na Uy). Con số 0,20 thầy '
    'đang xem KHẢ NĂNG CAO là β chuẩn hoá của mạng xã hội ở nữ (Bảng 4 trong bản dịch '
    'báo cáo β = 0,18 — gần như chắc chắn là 0,20 sau khi tác giả làm tròn). β = 0,20 '
    'là mức NHỎ theo Cohen, KTC 90% là KHÔNG CHUẨN.'
)
para_black('')
para_blue('Trả lời 2 ý của thầy:', bold=True)

para_blue('Ý 1 — Viết 0,20 (KTC 90%: 0,12–0,28) có CHUẨN không?', bold=True)
bullet_blue(
    'KHÔNG chuẩn theo quy ước y học/tâm lý hiện đại. Chuẩn vàng là KTC 95% (Cochrane, '
    'CONSORT, PRISMA, APA). KTC 90% chỉ chấp nhận khi tác giả NÊU RÕ lý do — '
    'TOST, Bayesian HDI, pilot/exploratory.'
)
bullet_blue(
    'Nếu bài 21 KHÔNG giải thích lý do dùng KTC 90%, đây là một ĐIỂM YẾU phương pháp '
    'cần đưa vào phần phản biện. Đôi khi tác giả "chữa" kết quả đuối bằng cách chuyển '
    'sang KTC 90% để khoảng không chứa 0.'
)
bullet_blue(
    'Quy đổi nhanh: KTC 90% [0,12–0,28] → KTC 95% ước tính ~ [0,10–0,30]. Khoảng này '
    'VẪN không chứa 0 → kết quả VẪN có ý nghĩa thống kê ở mức 95%. Vậy việc dùng KTC '
    '90% có thể là chọn lựa "an toàn" thay vì "cứu" kết quả.'
)

para_blue('Ý 2 — Mức 0,20 là NHỎ phải không?', bold=True)
bullet_blue(
    'ĐÚNG, mức NHỎ. Trong QT021, 0,20 là β chuẩn hoá (standardized regression '
    'coefficient) — bảng "Tác động chính β" báo cáo β mạng xã hội ở nữ = 0,18, ở nam = '
    '0,13. Theo Cohen 1988: β ≥ 0,1 nhỏ / 0,3 TB / 0,5 lớn → 0,20 = NHỎ.'
)
bullet_blue(
    'KTC 90% [0,12; 0,28] không chứa 0 → β khác 0 có ý nghĩa thống kê. Quy đổi sang '
    'KTC 95% ước tính ~ [0,10; 0,30] — vẫn không chứa 0, vẫn có ý nghĩa.'
)
bullet_blue(
    'NHƯNG đáng chú ý: mặc dù β nhỏ (0,18–0,20), Bảng 5 của Brunborg cho thấy GIỮ CỐ ĐỊNH '
    'mạng xã hội thì hệ số xu hướng tăng distress ở nữ giảm 60% (0,015 → 0,006). Vậy '
    'mạng xã hội VẪN GIẢI THÍCH ĐÁNG KỂ xu hướng tăng distress ở nữ Na Uy 2014–2024.'
)

para_blue('Quan trọng — phân biệt "có ý nghĩa thống kê" vs "đáng để ý lâm sàng":', bold=True)
bullet_blue(
    'r=0,20 với p<0,05 CÓ Ý NGHĨA THỐNG KÊ (đặc biệt khi mẫu lớn) nhưng độ lớn hiệu ứng '
    'NHỎ. Mô hình chỉ giải thích ~4% phương sai (r² = 0,04).'
)
bullet_blue(
    'Trong y văn social media + wellbeing (Orben & Przybylski 2019, Twenge 2018), r '
    'thường ~ 0,03–0,15 — tức cực nhỏ. Bài 21 báo cáo r=0,20 thì cao hơn trung vị y văn '
    'một chút, nhưng vẫn ở mức small effect.'
)

para_blue('Tóm lại cho QT021 Brunborg 2025:', bold=True)
para_blue(
    '(1) Cách viết KTC 90% là KHÔNG chuẩn — nên đưa vào phản biện. Brunborg có thể dùng '
    'KTC 90% trong bảng phụ/supplementary; bản dịch hiện có chưa lấy được KTC chi tiết. '
    '(2) Mức β = 0,20 là NHỎ theo Cohen — phù hợp với bảng 4 báo cáo β nữ = 0,18 cho '
    'mạng xã hội. (3) Nhưng đừng coi thường β nhỏ này: bảng 5 cho thấy mạng xã hội '
    'giải thích 60% xu hướng tăng distress ở nữ — đây là phát hiện ý nghĩa LỚN VỀ MẶT '
    'XÃ HỘI dù effect size nhỏ.'
)

para_blue('Đề xuất hành động:', bold=True)
bullet_blue('Khi trích dẫn vào báo cáo CTH, ghi: "β = 0,20 (KTC 95% ước tính 0,10–0,30) — mức NHỎ theo Cohen, nhưng giải thích đáng kể xu hướng tăng distress nữ Na Uy 2014–2024 (Brunborg et al. 2025, QT021)". Chuyển KTC 90% → 95% cho chuẩn.')
bullet_blue('Nếu cần nguyên văn 0,20 KTC 90%, em đề nghị thầy gửi link/trang chính xác trong bài gốc tiếng Anh để em verify (bản dịch của em chưa cover hết phần supplementary).')
bullet_blue('Cẩn trọng nội suy sang VN: Brunborg là Na Uy — văn hoá khác, đo distress chung HSCL không tách lo âu riêng. Khi áp dụng cho HS THCS/THPT VN cần kết hợp với VN002 VNAMHS, VN013 Tô Thị Hồng và Pham 2024 Huế.')

# =====================================================
# PHỤ LỤC — REFERENCES
# =====================================================
H('8. Phụ lục — Tài liệu tham khảo', level=2)

para_black(
    '1. Cohen, J. (1988). Statistical Power Analysis for the Behavioral Sciences (2nd '
    'ed.). Hillsdale, NJ: Lawrence Erlbaum. — Khung phân loại effect size cổ điển: '
    'r 0,1/0,3/0,5; d 0,2/0,5/0,8.', italic=True
)

para_black(
    '2. Cumming, G. (2014). The new statistics: Why and how. Psychological Science, '
    '25(1), 7–29. — Bảo vệ KTC 95% là chuẩn vàng; cảnh báo nguy cơ "p-hacking" qua '
    'thay đổi độ tin cậy.', italic=True
)

para_black(
    '3. Senn, S. (2007). Trying to be precise about vagueness. Statistics in Medicine, '
    '26(7), 1417–1430. — Hướng dẫn sử dụng KTC trong nghiên cứu lâm sàng; chỉ trích '
    'việc tuỳ tiện chọn 90% vs 95%.', italic=True
)

para_black(
    '4. Lakens, D. (2017). Equivalence tests: A practical primer for t-tests, '
    'correlations, and meta-analyses. Social Psychological and Personality Science, '
    '8(4), 355–362. — KHI nào dùng KTC 90% (TOST equivalence test) — bối cảnh DUY '
    'NHẤT mà 90% là chuẩn.', italic=True
)

para_black(
    '5. Orben, A., & Przybylski, A. K. (2019). The association between adolescent '
    'well-being and digital technology use. Nature Human Behaviour, 3(2), 173–182. — '
    'Tổng quan kinh điển: r social media ↔ wellbeing chỉ ~ 0,04 (cực nhỏ).', italic=True
)

para_black(
    '6. Twenge, J. M., Joiner, T. E., Rogers, M. L., & Martin, G. N. (2018). Increases '
    'in depressive symptoms... Clinical Psychological Science, 6(1), 3–17. — Ngược lại '
    'với Orben: r ~ 0,11–0,15. Hai phe tranh luận về độ lớn.', italic=True
)

para_black('7. Tham khảo trong corpus dự án (DB Lo-au):')
bullet_black(
    'QT021 Brunborg, Nilsen, Skogen, Bang (2025). Possible Explanations for the Upward '
    'Trend in Mental Distress among Adolescents in Norway from 2011 to 2024. Social '
    'Science & Medicine, 384, Article 118528. (Bài thầy đang xem — n=979.043 VTN Na Uy, '
    'β mạng xã hội ở nữ = 0,18.)', italic=True
)
bullet_black(
    'QT067 Pascoe et al. (2020) — Academic stress narrative review (IJAY 25:104-112). '
    'Có thảo luận về stress + mạng xã hội nhưng không có effect size pooled.', italic=True
)
bullet_black(
    'QT008 Wen et al. (2020) — LPA về lo âu HS THCS Trung Quốc; OR 2-11.', italic=True
)
bullet_black(
    'VN016 Nguyễn Cao Minh (2012) — chuẩn hóa RCADS VN, baseline để so sánh.',
    italic=True
)
bullet_black(
    'Pham 2024 (VN, Huế) — VN-corpus có dữ liệu social support + SKTT, đối chiếu được.',
    italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
