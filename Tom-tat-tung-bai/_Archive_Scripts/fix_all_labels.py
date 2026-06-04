# -*- coding: utf-8 -*-
"""Fix Bảng 1/2 labels cho tất cả tóm tắt thiếu"""
import os, sys, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '03_Ban-dich'))
from tao_dich_template import add_table, shade_cell, set_col_width
import docx
from docx.shared import Pt, RGBColor
from docx.enum.table import WD_TABLE_ALIGNMENT

OUT = os.path.dirname(os.path.abspath(__file__))

def get_text(doc):
    return chr(10).join([p.text for p in doc.paragraphs])

def add_label_before_table(doc, table_idx, label_text):
    """Insert a bold heading paragraph referencing a table by index"""
    # Find the table element in the document body
    table_elem = doc.tables[table_idx]._tbl
    # Create a new paragraph element
    p_elem = docx.oxml.OxmlElement('w:p')
    r_elem = docx.oxml.OxmlElement('w:r')
    rpr = docx.oxml.OxmlElement('w:rPr')
    b = docx.oxml.OxmlElement('w:b')
    rpr.append(b)
    sz = docx.oxml.OxmlElement('w:sz')
    sz.set(docx.oxml.ns.qn('w:val'), '22')  # 11pt
    rpr.append(sz)
    rFonts = docx.oxml.OxmlElement('w:rFonts')
    rFonts.set(docx.oxml.ns.qn('w:ascii'), 'Times New Roman')
    rFonts.set(docx.oxml.ns.qn('w:hAnsi'), 'Times New Roman')
    rpr.append(rFonts)
    r_elem.append(rpr)
    t_elem = docx.oxml.OxmlElement('w:t')
    t_elem.text = label_text
    r_elem.append(t_elem)
    p_elem.append(r_elem)
    # Insert before the table
    table_elem.addprevious(p_elem)

def fix_file(filename, table_labels=None):
    """Add Bảng 1, Bảng 2 labels before tables"""
    filepath = os.path.join(OUT, filename)
    doc = docx.Document(filepath)
    text = get_text(doc)
    n_tables = len(doc.tables)

    if 'Bảng 1' in text and 'Bảng 2' in text:
        return False  # Already OK

    if n_tables == 0:
        return False  # No tables to label

    if table_labels is None:
        table_labels = {0: 'Bảng 1. Kết quả chính', 1: 'Bảng 2. Yếu tố liên quan / So sánh'}

    for idx, label in sorted(table_labels.items()):
        if idx < n_tables:
            add_label_before_table(doc, idx, label)

    doc.save(filepath)
    # Verify
    doc2 = docx.Document(filepath)
    text2 = get_text(doc2)
    b1 = 'Bảng 1' in text2
    b2 = 'Bảng 2' in text2
    return b1 and b2

# ===== Define custom labels per file =====
custom_labels = {
    '00_Mẫu tóm tắt bài 1.docx': {0: 'Bảng 1. Kết quả chính', 1: 'Bảng 2. So sánh'},
    '02_Saikia_et_al_2023.docx': {0: 'Bảng 1. Tỷ lệ lo âu/trầm cảm', 1: 'Bảng 2. Yếu tố liên quan', 2: 'Bảng 3. So sánh liên bài'},
    '03_Mandaknalli_Malusare_2021.docx': {0: 'Bảng 1. Tỷ lệ GAD theo DASS-21', 1: 'Bảng 2. Yếu tố nhân khẩu', 2: 'Bảng 3. So sánh'},
    '04_NSCH_2020.docx': {0: 'Bảng 1. Xu hướng lo âu trẻ em Mỹ', 1: 'Bảng 2. Phân tầng theo nhóm', 2: 'Bảng 3. So sánh'},
    '05_Alharbi_et_al_2019.docx': {0: 'Bảng 1. Tỷ lệ lo âu HS Saudi Arabia', 1: 'Bảng 2. Yếu tố liên quan', 2: 'Bảng 3. So sánh'},
    '06_Nakie_et_al_2022.docx': {0: 'Bảng 1. Tỷ lệ lo âu Ethiopia', 1: 'Bảng 2. Yếu tố nguy cơ', 2: 'Bảng 3. Mô hình hồi quy', 3: 'Bảng 4. So sánh liên bài'},
    '07_Chen_et_al_2023.docx': {0: 'Bảng 1. Tỷ lệ lo âu HS Trung Quốc', 1: 'Bảng 2. Yếu tố liên quan'},
    '08_Wen_et_al_2020.docx': {0: 'Bảng 1. Phân tích LPA — Các hồ sơ lo âu', 1: 'Bảng 2. Yếu tố dự báo'},
    '09_Qiu_et_al_2022.docx': {0: 'Bảng 1. Kết quả chính — Phong cách nuôi dạy', 1: 'Bảng 2. Vai trò trung gian'},
    '10_Xu_et_al_2021.docx': {0: 'Bảng 1. Tỷ lệ lo âu VTN Trung Quốc', 1: 'Bảng 2. Yếu tố liên quan'},
    '11_Bhardwaj_et_al_2020.docx': {0: 'Bảng 1. Tỷ lệ DASS-21 Chandigarh', 1: 'Bảng 2. So sánh'},
    'QT25_EpiPsychSci_2025.docx': {0: 'Bảng 1. Trạng thái CMH 2018–2022', 1: 'Bảng 2. Flourishing theo giới'},
    'QT35_SocialAnxiety_7Countries.docx': {0: 'Bảng 1. Tỷ lệ SAD theo quốc gia', 1: 'Bảng 2. SAD theo phân nhóm'},
    'VN14_HoangTrungHoc_2025.docx': {0: 'Bảng 1. Tỷ lệ lo âu/stress VTN VN', 1: 'Bảng 2. Yếu tố ảnh hưởng'},
    'VN1_Hoa_2024.docx': {0: 'Bảng 1. Tỷ lệ lo âu GAD-7 HS THPT Hà Nội', 1: 'Bảng 2. Khác biệt giới tính'},
}

# ===== Files with only 1 table — need to add a 2nd table =====
need_2nd_table = {
    'QT21_Norway_2025.docx': {
        'header': ['Yếu tố', 'Đóng góp', 'Ghi chú'],
        'rows': [['Bất mãn trường học', 'GIẢI THÍCH CHÍNH', 'Loại bỏ xu hướng khi giữ cố định'],
                 ['Mạng xã hội', 'Giải thích MỘT PHẦN', 'b giảm 0,016→0,006'],
                 ['Tối ở nhà', 'Đóng góp nhỏ', ''],
                 ['Khó khăn tài chính', 'KHÔNG giải thích', 'Tỷ lệ GIẢM']],
    },
    'QT22_ScreenTime_2025.docx': {
        'header': ['Giới', 'Trầm cảm T1 (b)', 'Lo âu T1 (b)', 'Trầm cảm dọc (b)'],
        'rows': [['Nữ', 'Reference', 'Reference', 'Reference'],
                 ['Nam', '−2,58***', '−3,87***', '−2,22***'],
                 ['Đa dạng giới', '6,73***', '3,26***', '1,60***']],
    },
    'QT23_JAACAP_US_2024.docx': {
        'header': ['Nhóm tuổi', '2013 (%)', '2021 (%)', 'AOR', 'p'],
        'rows': [['<12 tuổi', '9,3', '14,6', '1,66', '<0,001'],
                 ['12–14 tuổi', '9,7', '21,1', '2,38', '<0,001'],
                 ['15–17 tuổi', '10,1', '25,4', '2,93', '<0,001']],
    },
    'QT24_WHO_Europe_2025.docx': {
        'header': ['Khuyến nghị', 'Nội dung', 'Ý nghĩa cho VN'],
        'rows': [['Trường học', 'Tích hợp SKTT vào giáo dục', 'Wen 2020: OR = 0,562'],
                 ['Dịch vụ', 'Phát triển cộng đồng', 'VN: 8,4% tiếp cận'],
                 ['MXH', 'Kiểm soát, giáo dục kỹ năng số', 'Norway 2025'],
                 ['Nghịch cảnh', 'Sàng lọc ACEs + can thiệp', 'Ngô Anh Vinh 2024']],
    },
    'VN2_VNAMHS_2022.docx': {
        'header': ['Chỉ số', 'Tỷ lệ', 'Ghi chú'],
        'rows': [['Rối loạn tâm thần (DISC-5)', '2,3%', 'Chẩn đoán — thấp hơn sàng lọc'],
                 ['Lo âu', 'Phổ biến nhất', ''],
                 ['Tiếp cận dịch vụ', '8,4%', 'Rất thấp'],
                 ['Nữ > Nam', 'Có', 'Phù hợp toàn cầu']],
    },
    'VN3_Pham_2024.docx': {
        'header': ['Yếu tố', 'Liên quan', 'Hướng'],
        'rows': [['Hỗ trợ xã hội', 'Lo âu + Trầm cảm', 'Bảo vệ (giảm)'],
                 ['Giới tính nữ', 'Lo âu', 'Nữ > Nam'],
                 ['Chất lượng chăm sóc', 'SKTT', 'Tích cực']],
    },
}

# ===== Process =====
fixed = 0
failed = 0

for f in sorted(os.listdir(OUT)):
    if not f.endswith('.docx'): continue
    if not (f[0].isdigit() or f.startswith('VN') or f.startswith('QT')): continue

    filepath = os.path.join(OUT, f)
    doc = docx.Document(filepath)
    text = get_text(doc)

    if 'Bảng 1' in text and 'Bảng 2' in text:
        continue  # Already OK

    n_tables = len(doc.tables)
    if n_tables == 0:
        continue

    # If only 1 table and need a 2nd
    if n_tables == 1 and f in need_2nd_table:
        info = need_2nd_table[f]
        t = doc.add_table(rows=1+len(info['rows']), cols=len(info['header']))
        t.style = 'Table Grid'
        t.alignment = WD_TABLE_ALIGNMENT.CENTER
        for ci, h in enumerate(info['header']):
            cell = t.rows[0].cells[ci]
            cell.text = h
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            shade_cell(cell, 'D9E2F3')
        for ri, row_data in enumerate(info['rows']):
            for ci, val in enumerate(row_data):
                cell = t.rows[ri+1].cells[ci]
                cell.text = val
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        doc.save(filepath)
        doc = docx.Document(filepath)  # Reload
        n_tables = len(doc.tables)

    # Now add labels
    labels = custom_labels.get(f, {0: 'Bảng 1. Kết quả chính', 1: 'Bảng 2. So sánh / Bổ sung'})

    for idx in sorted(labels.keys()):
        if idx < n_tables:
            add_label_before_table(doc, idx, labels[idx])

    doc.save(filepath)

    # Verify
    doc2 = docx.Document(filepath)
    text2 = get_text(doc2)
    if 'Bảng 1' in text2 and 'Bảng 2' in text2:
        fixed += 1
        print(f'  OK: {f}')
    else:
        failed += 1
        b1 = 'Bảng 1' in text2
        b2 = 'Bảng 2' in text2
        print(f'  FAIL: {f} (B1={b1}, B2={b2}, tables={len(doc2.tables)})')

print(f'\nFixed: {fixed}, Failed: {failed}')
