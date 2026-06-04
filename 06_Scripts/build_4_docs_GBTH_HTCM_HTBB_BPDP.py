"""MASTER SCRIPT — Build 4 doc dien giai yeu to nho con lai:
1. GBTH (gan bo truong hoc) — Bang 3.29-3.30
2. HTCM (ho tro cha me) — Bang 3.33-3.34
3. HTBB (ho tro ban be) — Bang 3.35-3.36
4. BPDP (bien phap doi pho) — Bang 3.43-3.45

ALL FACTS VERIFIED tu file v2 chuong 3.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)


def _new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
    style = d.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return d

def _H(d, text, level=2, color=NAVY):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def _para(d, text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def _blue(d, text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def _table(d, header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def _cap(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def _ref(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

def _title(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY

# ===========================================================
# DOC 1: GBTH (gắn bó trường học) — chứa giao tiếp GV-HS
# ===========================================================
def build_GBTH():
    d = _new_doc()
    _title(d, 'DIỄN GIẢI YẾU TỐ GẮN BÓ TRƯỜNG HỌC (GBTH)\n— Bao gồm quan hệ GV-HS, bạn bè, môi trường —\nBảng 3.22, 3.29, 3.30 chương 3 luận án')
    _para(d, '')

    _H(d, '1. Khung khái niệm', 2)
    _para(d,
        'GẮN BÓ TRƯỜNG HỌC (school connectedness / school '
        'belonging) là cảm nhận của học sinh về việc THUỘC VỀ '
        'môi trường trường học — bao gồm mối quan hệ với giáo '
        'viên, bạn bè, và không gian trường. Đo bằng thang '
        'Psychological Sense of School Membership (PSSM; Goodenow, '
        '1993) — gồm các mục về cảm nhận thân thiện, được quan '
        'tâm, được tham gia.', indent=False
    )

    _H(d, '2. Diễn giải Bảng 3.22 — Mức độ và biểu hiện GBTH', 2)
    _cap(d, 'Bảng 1. Biểu hiện GBTH (rút từ Bảng 3.22)')
    _table(d, ['Mục', 'Nội dung', 'ĐTB', 'ĐLC', 'Thứ bậc'],
        [
            ['F. Gắn bó với trường học', 'Tổng nhóm', '52,60', '20,021', '4'],
            ['PSSM.10', 'Tham gia nhiều hoạt động ở trường', '57,69', '32,242', '1'],
            ['PSSM.8', 'Mọi người ở trường thân thiện với tôi', '57,40', '29,350', '2'],
            ['PSSM.13', 'Có thể là chính mình khi ở trường', '53,25', '33,580', '3'],
            ['PSSM.15', 'Mọi người ở trường biết tôi có thể đóng góp', '52,02', '30,687', '4'],
            ['PSSM.1', 'Cảm thấy là một phần của ngôi trường', '47,86', '32,273', '5'],
            ['PSSM.5', 'Phần lớn giáo viên quan tâm đến tôi', '47,36', '29,127', '6'],
        ]
    )
    _para(d, '')
    _para(d,
        'ĐTB tổng = 52,60 (ĐLC = 20,021) — TRÊN GIỮA THANG. Thứ '
        'bậc 4 trong nhóm yếu tố bảo vệ. Trong sáu mục, '
        '"giáo viên quan tâm" có ĐTB THẤP NHẤT (47,36) — cảnh báo '
        'rằng QUAN HỆ GV-HS là điểm yếu của gắn bó trường học. '
        '"Tham gia hoạt động" có ĐTB CAO NHẤT (57,69) — học sinh '
        'cảm nhận tham gia tốt hơn cảm nhận quan hệ với GV.'
    )

    _H(d, '3. Diễn giải Bảng 3.29 — Fit indices SEM GBTH', 2)
    _cap(d, 'Bảng 2. Fit indices các mô hình GBTH → RLLA (Bảng 3.29)')
    _table(d, ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI'],
        [
            ['GBTH–RLLATQ', '198,365 (76)', '2,610', '0,970', '0,964', '0,035', '0,029–0,040'],
            ['GBTH–RLLACL', '112,906 (43)', '2,626', '0,973', '0,965', '0,035', '0,027–0,043'],
            ['GBTH–RLLAXH', '123,336 (43)', '2,868', '0,971', '0,963', '0,037', '0,030–0,045'],
            ['GBTH–RLLA', '112,352 (34)', '3,304', '0,970', '0,960', '0,041', '0,033–0,050'],
            ['GBTH–RLLA (2 factor)', '79,072 (26)', '3,041', '0,978', '0,970', '0,039', '0,029–0,049'],
        ]
    )
    _para(d, '')
    _para(d,
        'Tất cả 5 mô hình đạt fit TỐT theo Hu & Bentler (1999): '
        'CFI 0,970–0,978; RMSEA 0,035–0,041 (đều dưới 0,05); '
        'χ²/df 2,610–3,304 (gần ngưỡng lý tưởng 3). Mô hình 2 '
        'factor có CFI cao nhất (0,978).'
    )

    _H(d, '4. Diễn giải Bảng 3.30 — Hệ số tác động GBTH → RLLA', 2)
    _cap(d, 'Bảng 3. Hệ số tác động GBTH → RLLA (Bảng 3.30)')
    _table(d, ['Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
        [
            ['GBTH → RLLATQ', '−0,108', '0,049', '−3,073', '0,002', '—'],
            ['GBTH → RLLACL', '0,014', '0,052', '0,390', '0,696', '— (NS)'],
            ['GBTH → RLLAXH', '−0,187', '0,054', '−5,005', '< 0,001', '—'],
            ['GBTH → RLLA', '−0,155', '0,255', '−4,326', '< 0,001', '0,024'],
            ['GBTH → RLLA (2 factor)', '−0,146', '0,278', '−2,648', '0,008', '0,021'],
        ]
    )
    _para(d, '')
    _H(d, 'Diễn giải:', 3)
    _para(d,
        'GBTH → RLLATQ: β = −0,108 (p = 0,002) — tác động BẢO VỆ '
        'NHỎ. Học sinh có cảm nhận gắn bó trường tốt hơn → ít lo '
        'âu lan tỏa hơn.'
    )
    _para(d,
        'GBTH → RLLACL: β = 0,014 (p = 0,696) — KHÔNG có ý nghĩa '
        'thống kê. Gắn bó trường KHÔNG bảo vệ trước lo âu chia '
        'ly. Phù hợp với khung phát triển: lo âu chia ly liên '
        'quan đến gắn bó với CHA MẸ chứ không phải với trường '
        'học.'
    )
    _para(d,
        'GBTH → RLLAXH: β = −0,187 (p < 0,001) — tác động BẢO VỆ '
        'MẠNH NHẤT trong ba dạng. Phát hiện logic: gắn bó với '
        'trường (gồm bạn bè, giáo viên) GIẢM lo âu xã hội. Đây '
        'là yếu tố BẢO VỆ ĐẶC BIỆT cho RLLAXH.'
    )
    _para(d,
        'GBTH → RLLA tổng: β = −0,155; R² = 2,4% — tác động '
        'tổng NHỎ. Nói cách khác, gắn bó trường có tác động '
        'BẢO VỆ ý nghĩa thống kê nhưng cường độ KHIÊM TỐN.'
    )

    _H(d, '5. Phát hiện đặc biệt', 2)
    _para(d,
        'Pattern KHÁC với các yếu tố bảo vệ khác: GBTH bảo vệ '
        'MẠNH NHẤT cho RLLAXH (lo âu xã hội), KHÔNG có tác '
        'động cho RLLAC (chia ly). Ngược với HTCM (hỗ trợ '
        'cha mẹ — sẽ phân tích trong doc riêng) cũng có '
        'pattern tương tự nhưng qua cơ chế khác.'
    )
    _para(d,
        'Điểm yếu của gắn bó trường học VN: ĐTB của "giáo '
        'viên quan tâm" (PSSM.5 = 47,36) THẤP NHẤT trong sáu '
        'mục — cảnh báo về QUAN HỆ GV-HS. Cần can thiệp đào '
        'tạo giáo viên về kỹ năng quan tâm + lắng nghe học '
        'sinh.'
    )

    _H(d, '6. CÂU TRẢ LỜI', 2)
    _blue(d, 'Diễn giải GBTH (gắn bó trường học):', bold=True)
    _blue(d,
        '(1) ĐTB = 52,60 (ĐLC = 20,021) — trên giữa thang. '
        '"Giáo viên quan tâm" (PSSM.5 = 47,36) THẤP NHẤT — '
        'điểm yếu về quan hệ GV-HS.'
    )
    _blue(d,
        '(2) Fit indices TỐT: CFI 0,970–0,978; RMSEA '
        '0,035–0,041 (5 mô hình).'
    )
    _blue(d,
        '(3) β GBTH → RLLATQ = −0,108 (p = 0,002) NHỎ; β '
        '→ RLLACL = 0,014 (KHÔNG sig); β → RLLAXH = '
        '−0,187 (p < 0,001) MẠNH NHẤT. R² tổng = 2,4% — '
        'tác động bảo vệ KHIÊM TỐN.'
    )
    _blue(d,
        '(4) Pattern: GBTH bảo vệ ĐẶC BIỆT cho lo âu xã '
        'hội (logic — qua bạn bè + GV), KHÔNG có tác động '
        'cho lo âu chia ly (logic — chia ly liên quan cha '
        'mẹ).'
    )

    _H(d, '7. Tài liệu tham khảo', 2)
    for r in [
        'Goodenow, C. (1993). The psychological sense of school membership among adolescents: Scale development and educational correlates. Psychology in the Schools, 30(1), 79–90.',
        'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
    ]:
        _ref(d, r)

    out = OUT_DIR / 'Dien_giai_yeu_to_GBTH_gan_bo_truong_hoc.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')


# ===========================================================
# DOC 2: HTCM (hỗ trợ cha mẹ)
# ===========================================================
def build_HTCM():
    d = _new_doc()
    _title(d, 'DIỄN GIẢI YẾU TỐ HỖ TRỢ TỪ CHA MẸ (HTCM)\n— Bảng 3.22 (mục G), 3.33, 3.34 chương 3 luận án —')
    _para(d, '')

    _H(d, '1. Khung khái niệm', 2)
    _para(d,
        'HỖ TRỢ TỪ CHA MẸ là cảm nhận của học sinh về sự sẵn '
        'sàng giúp đỡ về CẢM XÚC, THÔNG TIN, và VẬT CHẤT từ gia '
        'đình. Đo bằng tiểu thang Family thuộc Multidimensional '
        'Scale of Perceived Social Support (MSPSS; Zimet và '
        'cộng sự, 1988).', indent=False
    )

    _H(d, '2. Diễn giải Bảng 3.22 — Mức độ HTCM', 2)
    _cap(d, 'Bảng 1. Biểu hiện HTCM (rút từ Bảng 3.22 mục G)')
    _table(d, ['Mục', 'Nội dung', 'ĐTB', 'ĐLC', 'Thứ bậc'],
        [
            ['G. Hỗ trợ từ cha mẹ', 'Tổng nhóm', '57,65', '27,001', '1'],
            ['MSPSS.3', 'Gia đình thực sự cố gắng giúp đỡ', '65,48', '31,441', '1'],
            ['MSPSS.4', 'Hỗ trợ về cảm xúc từ gia đình', '59,82', '32,806', '2'],
            ['MSPSS.11', 'Gia đình giúp đưa ra quyết định', '57,75', '31,969', '3'],
            ['MSPSS.8', 'Có thể nói chuyện về vấn đề', '47,54', '34,157', '4'],
        ]
    )
    _para(d, '')
    _para(d,
        'ĐTB tổng = 57,65 — CAO NHẤT trong nhóm yếu tố bảo vệ '
        '(thứ bậc 1). MSPSS.3 (gia đình giúp đỡ) có ĐTB cao '
        'nhất 65,48 — gần 2/3 thang đo. Điểm yếu: MSPSS.8 '
        '(nói chuyện về vấn đề) có ĐTB chỉ 47,54 — gợi ý '
        'CHA MẸ Việt Nam giỏi HỖ TRỢ THỰC TẾ nhưng yếu hơn '
        'về GIAO TIẾP CẢM XÚC.'
    )

    _H(d, '3. Diễn giải Bảng 3.33 — Fit indices SEM HTCM', 2)
    _cap(d, 'Bảng 2. Fit indices HTCM → RLLA (Bảng 3.33)')
    _table(d, ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI'],
        [
            ['HTCM–RLLATQ', '136,234 (43)', '3,168', '0,981', '0,975', '0,040', '0,033–0,048'],
            ['HTCM–RLLACL', '87,436 (19)', '4,602', '0,980', '0,970', '0,052', '0,041–0,063'],
            ['HTCM–RLLAXH', '61,229 (19)', '3,223', '0,988', '0,983', '0,041', '0,029–0,052'],
            ['HTCM–RLLA', '84,551 (13)', '6,504', '0,979', '0,966', '0,064', '0,051–0,077'],
            ['HTCM–RLLA (2 factor)', '48,008 (8)', '6,001', '0,988', '0,977', '0,061', '0,045–0,078'],
        ]
    )
    _para(d, '')
    _para(d,
        'Fit TỐT đến XUẤT SẮC: CFI 0,979–0,988 (đều > 0,95); '
        'RMSEA 0,040–0,064 (4/5 dưới 0,06; mô hình tổng vượt '
        'nhẹ). χ²/df mô hình tổng và 2 factor cao (6,5 và '
        '6,0) — phản ánh độ phức tạp khi tổng hợp 3 dạng RLLA.'
    )

    _H(d, '4. Diễn giải Bảng 3.34 — Hệ số tác động HTCM → RLLA', 2)
    _cap(d, 'Bảng 3. Hệ số tác động HTCM → RLLA (Bảng 3.34)')
    _table(d, ['Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
        [
            ['HTCM → RLLATQ', '−0,172', '0,021', '−5,299', '< 0,001', '—'],
            ['HTCM → RLLACL', '0,000', '0,022', '−0,003', '0,997', '— (NS)'],
            ['HTCM → RLLAXH', '−0,273', '0,022', '−7,896', '< 0,001', '—'],
            ['HTCM → RLLA', '−0,234', '0,123', '−7,035', '< 0,001', '0,055'],
            ['HTCM → RLLA (2 factor)', '−0,222', '0,137', '−4,902', '< 0,001', '0,049'],
        ]
    )
    _para(d, '')
    _H(d, 'Diễn giải:', 3)
    _para(d,
        'HTCM → RLLATQ: β = −0,172 (p < 0,001) — tác động bảo '
        'vệ NHỎ–TRUNG BÌNH. Hỗ trợ cha mẹ tốt → ít lo âu lan '
        'tỏa hơn.'
    )
    _para(d,
        'HTCM → RLLACL: β = 0,000 (p = 0,997) — TUYỆT ĐỐI '
        'KHÔNG có tác động! Đây là phát hiện ĐẶC BIỆT — TRÁI '
        'với kỳ vọng thông thường rằng "có hỗ trợ cha mẹ tốt '
        'thì giảm lo âu chia ly". Có thể giải thích: học '
        'sinh THCS Việt Nam dù có hỗ trợ cha mẹ tốt vẫn có '
        'thể mắc lo âu chia ly do đặc thù gắn bó văn hóa Á '
        '(không tách khỏi gia đình sớm như phương Tây).'
    )
    _para(d,
        'HTCM → RLLAXH: β = −0,273 (p < 0,001) — tác động '
        'BẢO VỆ MẠNH NHẤT trong ba dạng. Cha mẹ hỗ trợ tốt '
        '→ giảm lo âu xã hội. Cơ chế: cha mẹ là người tạo '
        '"nền tảng an toàn" cho học sinh thử nghiệm tương '
        'tác xã hội.'
    )
    _para(d,
        'HTCM → RLLA tổng: β = −0,234; R² = 5,5% — cường '
        'độ TRUNG BÌNH–CAO so với các YTBV khác (lớn hơn '
        'GBTH với β = −0,155).'
    )

    _H(d, '5. Phát hiện đặc biệt', 2)
    _para(d,
        'Phát hiện CRITICAL: HTCM → RLLACL có β = 0,000 — '
        'KHÔNG có tác động bảo vệ với lo âu chia ly. Pattern '
        'này TRÁI với kỳ vọng và TRÁI với pattern phương '
        'Tây. Cần phản biện riêng: có thể HTCM ở Việt Nam '
        'củng cố sự gắn bó với cha mẹ → khó tách khỏi cha '
        'mẹ → KHÔNG giảm lo âu chia ly mà còn có thể duy '
        'trì.'
    )
    _para(d,
        'Hỗ trợ cảm xúc < hỗ trợ thực tế: MSPSS.8 (nói '
        'chuyện về vấn đề) có ĐTB chỉ 47,54 trong khi các '
        'mục về giúp đỡ thực tế có ĐTB > 57. Cha mẹ Việt '
        'Nam giỏi HỖ TRỢ THỰC TẾ nhưng yếu về GIAO TIẾP '
        'CẢM XÚC. Phù hợp phát hiện Phạm 2024 (VN003): '
        'chăm sóc cảm xúc β = −0,40 mạnh hơn chăm sóc '
        'thể chất.'
    )

    _H(d, '6. CÂU TRẢ LỜI', 2)
    _blue(d, 'Diễn giải HTCM (hỗ trợ cha mẹ):', bold=True)
    _blue(d,
        '(1) ĐTB = 57,65 — CAO NHẤT trong YTBV. Gia đình '
        'giúp đỡ (MSPSS.3) cao nhất 65,48; nói chuyện về '
        'vấn đề (MSPSS.8) thấp nhất 47,54.'
    )
    _blue(d,
        '(2) Fit TỐT: CFI 0,979–0,988; RMSEA 0,040–0,064.'
    )
    _blue(d,
        '(3) β HTCM → RLLATQ = −0,172 (p < 0,001) NHỎ; β '
        '→ RLLACL = 0,000 (p = 0,997 — KHÔNG có tác '
        'động!); β → RLLAXH = −0,273 (p < 0,001) MẠNH '
        'NHẤT. R² tổng = 5,5%.'
    )
    _blue(d,
        '(4) Phát hiện ĐẶC BIỆT: HTCM KHÔNG có tác động '
        'cho lo âu chia ly — có thể do văn hóa Á gắn bó '
        'gia đình mạnh.'
    )

    _H(d, '7. Tài liệu tham khảo', 2)
    for r in [
        'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
        'Pham, S. T., Duong, T. T. T., Nguyen, H. P. T., & Truong, X. N. T. (2024). The correlation between quality of care and mental health and behavioral problems among Vietnamese adolescents in social support facilities. [VN003 trong cơ sở dữ liệu dự án.]',
        'Zimet, G. D., Dahlem, N. W., Zimet, S. G., & Farley, G. K. (1988). The Multidimensional Scale of Perceived Social Support. Journal of Personality Assessment, 52(1), 30–41.',
    ]:
        _ref(d, r)

    out = OUT_DIR / 'Dien_giai_yeu_to_HTCM_ho_tro_cha_me.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')


# ===========================================================
# DOC 3: HTBB (hỗ trợ bạn bè)
# ===========================================================
def build_HTBB():
    d = _new_doc()
    _title(d, 'DIỄN GIẢI YẾU TỐ HỖ TRỢ TỪ BẠN BÈ (HTBB)\n— Phát hiện ĐẶC BIỆT: GẦN NHƯ KHÔNG có tác động —\nBảng 3.22 (mục H), 3.35, 3.36 chương 3 luận án')
    _para(d, '')

    _H(d, '1. Khung khái niệm', 2)
    _para(d,
        'HỖ TRỢ TỪ BẠN BÈ đo bằng tiểu thang Friends thuộc '
        'MSPSS (Zimet và cộng sự, 1988). Theo y văn quốc tế, '
        'hỗ trợ bạn bè được kỳ vọng là yếu tố BẢO VỆ MẠNH cho '
        'sức khỏe tâm thần thanh thiếu niên.', indent=False
    )

    _H(d, '2. Diễn giải Bảng 3.36 — Hệ số tác động HTBB → RLLA', 2)
    _cap(d, 'Bảng 1. Hệ số tác động HTBB → RLLA (Bảng 3.36)')
    _table(d, ['Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
        [
            ['HTBB → RLLATQ', '−0,015', '0,021', '−0,459', '0,646', '— (NS)'],
            ['HTBB → RLLACL', '−0,019', '0,023', '0,558', '0,577', '— (NS)'],
            ['HTBB → RLLAXH', '−0,079', '0,023', '−2,318', '0,020', '—'],
            ['HTBB → RLLA', '−0,044', '0,133', '−1,331', '0,183', '0,002 (NS)'],
            ['HTBB → RLLA (2 factor) [LỖI]', '−0,039', '0,144', '−0,438', '0,661', '0,002'],
        ]
    )
    _para(d, '')

    _H(d, 'Phát hiện ĐẶC BIỆT — TRÁI VỚI KỲ VỌNG', 2)
    _para(d,
        'HTBB → RLLATQ: β = −0,015 (p = 0,646) — KHÔNG có ý '
        'nghĩa thống kê.'
    )
    _para(d,
        'HTBB → RLLACL: β = −0,019 (p = 0,577) — KHÔNG có ý '
        'nghĩa thống kê.'
    )
    _para(d,
        'HTBB → RLLAXH: β = −0,079 (p = 0,020) — có ý nghĩa '
        'thống kê NHƯNG cường độ RẤT NHỎ (|β| < 0,10).'
    )
    _para(d,
        'HTBB → RLLA tổng: β = −0,044 (p = 0,183) — KHÔNG có '
        'ý nghĩa thống kê. R² = 0,2% — TỐI THIỂU.'
    )
    _para(d,
        'Đây là phát hiện CRITICAL: hỗ trợ bạn bè GẦN NHƯ '
        'KHÔNG có tác động bảo vệ trước RLLA ở học sinh THCS '
        'Việt Nam — TRÁI với y văn quốc tế. Có 4 khả năng '
        'giải thích.'
    )

    _H(d, '3. Bốn khả năng giải thích', 2)
    _para(d,
        'KHẢ NĂNG 1 — Đặc thù lứa tuổi THCS. Học sinh lớp '
        '6–9 (11–15 tuổi) có thể CHƯA phát triển đầy đủ kỹ '
        'năng dựa vào bạn bè. Hỗ trợ bạn bè trở nên quan '
        'trọng hơn ở THPT (15–18 tuổi).'
    )
    _para(d,
        'KHẢ NĂNG 2 — Văn hóa Á tập trung vào gia đình. '
        'Trong văn hóa Việt Nam, gia đình (cha mẹ) đóng vai '
        'trò TRUNG TÂM trong hỗ trợ — bạn bè ở vị trí THỨ '
        'YẾU. Phù hợp với HTCM β = −0,234 mạnh hơn HTBB '
        'rất nhiều.'
    )
    _para(d,
        'KHẢ NĂNG 3 — Chất lượng vs số lượng. Học sinh có '
        'thể có nhiều bạn nhưng quan hệ NÔNG. Phù hợp với '
        'phát hiện Fassi 2025 (Nature Hum Behav, QT027): '
        'thanh thiếu niên có nhiều bạn online nhưng kém '
        'hài lòng. Cần đo CHẤT LƯỢNG quan hệ thay vì chỉ '
        'cảm nhận hỗ trợ.'
    )
    _para(d,
        'KHẢ NĂNG 4 — Mô hình SEM lỗi. Mô hình HTBB → RLLA '
        '(2 factor) được tác giả luận án đánh dấu là "mô '
        'hình lỗi" — có thể có vấn đề specification. Cần '
        'thử cấu trúc khác.'
    )

    _H(d, '4. Hàm ý', 2)
    _para(d,
        'HÀM Ý 1 — KHÔNG dựa vào "tăng hỗ trợ bạn bè" làm '
        'can thiệp chính cho RLLA học sinh THCS Việt Nam. '
        'Hiệu quả không có bằng chứng tại lứa tuổi này.'
    )
    _para(d,
        'HÀM Ý 2 — Tập trung vào CHẤT LƯỢNG quan hệ. '
        'Chương trình peer support (Murphy 2024 — QT066) '
        'có thể vẫn có giá trị nhưng cần CẤU TRÚC RÕ '
        'RÀNG về vai trò peer worker.'
    )
    _para(d,
        'HÀM Ý 3 — Lặp lại nghiên cứu ở THPT. Kiểm tra '
        'xem hỗ trợ bạn bè có tác động bảo vệ ở lứa tuổi '
        'cao hơn (15–18) không.'
    )

    _H(d, '5. CÂU TRẢ LỜI', 2)
    _blue(d, 'Diễn giải HTBB (hỗ trợ bạn bè) — phát hiện CRITICAL:', bold=True)
    _blue(d,
        '(1) HTBB GẦN NHƯ KHÔNG có tác động bảo vệ trước '
        'RLLA ở HS THCS Việt Nam. β → RLLATQ = −0,015 (NS); '
        'β → RLLACL = −0,019 (NS); β → RLLAXH = −0,079 (sig '
        'nhưng |β| < 0,10).'
    )
    _blue(d,
        '(2) β → RLLA tổng = −0,044 (p = 0,183 NS); R² = '
        '0,2% — TRÁI với y văn quốc tế.'
    )
    _blue(d,
        '(3) Bốn khả năng: tuổi THCS quá sớm, văn hóa Á '
        'tập trung gia đình, chất lượng quan hệ thấp, mô '
        'hình SEM có vấn đề specification.'
    )
    _blue(d,
        '(4) Hàm ý: KHÔNG dựa vào hỗ trợ bạn bè làm can '
        'thiệp chính. Cần đo CHẤT LƯỢNG quan hệ + lặp '
        'lại ở THPT.'
    )

    _H(d, '6. Tài liệu tham khảo', 2)
    for r in [
        'Fassi, L., và cộng sự. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9(6), 1283–1299. [QT027 trong cơ sở dữ liệu.]',
        'Murphy, R., và cộng sự. (2024). A systematic scoping review of peer support interventions in integrated primary youth mental health care. Journal of Community Psychology, 52(1), 154–180. [QT066.]',
        'Zimet, G. D., Dahlem, N. W., Zimet, S. G., & Farley, G. K. (1988). The Multidimensional Scale of Perceived Social Support. Journal of Personality Assessment, 52(1), 30–41.',
    ]:
        _ref(d, r)

    out = OUT_DIR / 'Dien_giai_yeu_to_HTBB_ho_tro_ban_be.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')


# ===========================================================
# DOC 4: BPĐP (biện pháp đối phó)
# ===========================================================
def build_BPDP():
    d = _new_doc()
    _title(d, 'DIỄN GIẢI YẾU TỐ BIỆN PHÁP ĐỐI PHÓ (BPĐP)\n— Phát hiện CRITICAL: β = 0,749 nhưng fit KÉM —\nBảng 3.43, 3.44, 3.45 chương 3 luận án')
    _para(d, '')

    _H(d, '1. Khung khái niệm', 2)
    _para(d,
        'BIỆN PHÁP ĐỐI PHÓ (coping strategies) — theo Carver '
        '(1997) Brief-COPE — phân thành ba nhóm chính: '
        '(a) Tránh né vấn đề (avoidance); (b) Tập trung giải '
        'quyết vấn đề (problem-focused); (c) Tìm kiếm sự hỗ '
        'trợ (support-seeking). Brief-COPE 28 mục, 14 nhân tố '
        '(Carver 1997).', indent=False
    )

    _H(d, '2. Diễn giải Bảng 3.43 — Mức độ và biểu hiện BPĐP', 2)
    _cap(d, 'Bảng 1. Mức độ ba nhóm BPĐP (rút từ Bảng 3.43)')
    _table(d, ['Nhóm', 'Đặc trưng', 'ĐTB', 'ĐLC', 'Thứ bậc'],
        [
            ['Tìm kiếm sự hỗ trợ', 'Hỏi giúp, chia sẻ', '55,00', '24,288', '1'],
            ['Tập trung giải quyết vấn đề', 'Lập kế hoạch, hành động', '53,18', '19,974', '2'],
            ['Tránh né vấn đề', 'Tự trách, từ bỏ, không tin', '41,34', '22,641', '3'],
        ]
    )
    _para(d, '')
    _para(d,
        'Học sinh THCS Việt Nam dùng nhiều nhất là TÌM KIẾM '
        'HỖ TRỢ (ĐTB = 55,00). Tránh né có ĐTB thấp nhất '
        '(41,34) — gợi ý đa số học sinh KHÔNG có xu hướng '
        'tránh né. Tuy nhiên, đoạn dưới đây sẽ cho thấy '
        'pattern phức tạp khi đi vào SEM.'
    )

    _H(d, '3. Diễn giải Bảng 3.44 — Fit indices SEM BPĐP', 2)
    _cap(d, 'Bảng 2. Fit indices BPĐP → RLLA (Bảng 3.44)')
    _table(d, ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI'],
        [
            ['BPĐP–RLLATQ', '327,439 (34)', '9,631', '0,911', '0,882', '0,080', '0,072–0,088'],
            ['BPĐP–RLLACL', '186,218 (13)', '14,324', '0,883', '0,811', '0,099', '0,087–0,112'],
            ['BPĐP–RLLAXH', '277,710 (13)', '21,362', '0,865', '0,782', '0,123', '0,110–0,136'],
            ['BPĐP–RLLA (3 factors)', '234,071 (8)', '29,259', '0,882', '0,779', '0,145', '0,129–0,161'],
            ['BPĐP–RLLA (2 factor)', '229,057 (4)', '57,264', '0,871', '0,678', '0,204', '0,182–0,227'],
        ]
    )
    _para(d, '')
    _para(d,
        'Fit indices KÉM trên TẤT CẢ mô hình: CFI 0,865–0,911 '
        '(dưới ngưỡng 0,95); TLI 0,678–0,882 (dưới 0,95); '
        'RMSEA 0,080–0,204 (vượt ngưỡng 0,06; mô hình tổng '
        'lên đến 0,204 = RẤT KÉM); χ²/df 9,6–57,3 (vượt xa '
        'ngưỡng 5). Đây là mô hình có FIT KÉM NHẤT trong '
        'toàn bộ chương 3.'
    )
    _para(d,
        'Lý do có thể: Brief-COPE có 14 nhân tố — gộp thành '
        '3 nhóm có thể MẤT CẤU TRÚC. Tác giả luận án đã '
        'thẳng thắn báo cáo và đề xuất phân tách lại theo '
        '14 nhân tố trong nghiên cứu tiếp theo.'
    )

    _H(d, '4. Diễn giải Bảng 3.45 — Hệ số tác động BPĐP → RLLA', 2)
    _cap(d, 'Bảng 3. Hệ số tác động BPĐP → RLLA (Bảng 3.45)')
    _table(d, ['Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
        [
            ['BPĐP → RLLATQ', '0,749', '0,025', '6,488', '< 0,001', '0,561'],
            ['BPĐP → RLLACL', '0,132', '0,040', '2,912', '0,004', '0,017'],
            ['BPĐP → RLLAXH', '0,670', '0,027', '4,920', '< 0,001', '0,449'],
            ['BPĐP → RLLA (3 factors)', '0,735', '0,156', '6,137', '< 0,001', '0,540'],
            ['BPĐP → RLLA (2 factor)', '0,728', '0,157', '6,021', '< 0,001', '0,529'],
        ]
    )
    _para(d, '')

    _H(d, 'Phát hiện CRITICAL — β CỰC LỚN nhưng dấu DƯƠNG', 2)
    _para(d,
        'BPĐP → RLLATQ: β = 0,749 — đây là CƯỜNG ĐỘ TÁC ĐỘNG '
        'LỚN NHẤT trong toàn bộ chương 3 luận án — vượt cả β '
        'ALHT = 0,510 và |β TTr| = 0,455. R² = 56,1% — biện '
        'pháp đối phó MỘT MÌNH giải thích hơn một nửa phương '
        'sai RLLATQ.'
    )
    _para(d,
        'BPĐP → RLLAXH: β = 0,670 — CƯỜNG ĐỘ RẤT LỚN. '
        'R² = 44,9%. Pattern tương tự RLLATQ.'
    )
    _para(d,
        'BPĐP → RLLACL: β = 0,132 — yếu (p = 0,004). Lo âu '
        'chia ly ít liên quan với BPĐP.'
    )
    _para(d,
        'BPĐP → RLLA tổng: β = 0,735 (3 factors) hoặc 0,728 '
        '(2 factor); R² = 0,540 hoặc 0,529 — ~53% phương sai '
        'RLLA tổng giải thích bởi BPĐP.'
    )
    _para(d,
        'DẤU β DƯƠNG là phát hiện ĐẶC BIỆT NGHỊCH LÝ — TRÁI '
        'với giả định trực giác "đối phó nhiều thì lo âu '
        'giảm". Phù hợp với hiện tượng MALADAPTIVE COPING '
        'ESCALATION (Compas và cộng sự, 2017): học sinh càng '
        'lo âu càng dùng nhiều biện pháp đối phó, nhưng nếu '
        'không hiệu quả thì lo âu vẫn duy trì hoặc tăng.'
    )

    _H(d, '5. Cảnh báo về diễn giải', 2)
    _para(d,
        'KẾT HỢP β LỚN + FIT KÉM tạo ra cảnh báo nghiêm '
        'trọng:'
    )
    _para(d,
        '• β = 0,749 hấp dẫn nhưng KHÔNG đáng tin do fit '
        'indices KÉM (RMSEA 0,080–0,204; CFI 0,865–0,911).\n'
        '• Mô hình có thể chưa specify đúng — gộp 3 nhóm BPĐP '
        'thành 1 biến tiềm ẩn không bắt được cấu trúc thực.\n'
        '• Cần phân tách lại theo Brief-COPE 14 nhân tố '
        '(Carver 1997).\n'
        '• KHÔNG NÊN báo cáo β = 0,749 làm phát hiện chính '
        'trong báo cáo — chỉ trình bày như phát hiện EXPLORATORY.', indent=False
    )

    _H(d, '6. CÂU TRẢ LỜI', 2)
    _blue(d, 'Diễn giải BPĐP — phát hiện CRITICAL:', bold=True)
    _blue(d,
        '(1) Mức độ: ĐTB ba nhóm — Tìm hỗ trợ 55,00 > Giải '
        'quyết 53,18 > Tránh né 41,34. HS THCS VN dùng nhiều '
        'nhất tìm hỗ trợ.'
    )
    _blue(d,
        '(2) Fit indices KÉM TRÊN TẤT CẢ mô hình: CFI '
        '0,865–0,911 (dưới 0,95); RMSEA 0,080–0,204 (vượt '
        'xa 0,06). Mô hình có vấn đề specification.'
    )
    _blue(d,
        '(3) β CỰC LỚN: BPĐP → RLLATQ = 0,749 (R² = '
        '56,1%); → RLLAXH = 0,670 (R² = 44,9%); → RLLACL '
        '= 0,132 (yếu); → RLLA tổng = 0,735 (R² = 54,0%). '
        'Lớn nhất toàn chương 3.'
    )
    _blue(d,
        '(4) DẤU β DƯƠNG nghịch lý — phù hợp '
        'MALADAPTIVE COPING ESCALATION (Compas 2017): '
        'lo âu càng cao càng dùng nhiều BPĐP nhưng '
        'không hiệu quả.'
    )
    _blue(d,
        '(5) Cảnh báo: KHÔNG báo cáo β = 0,749 làm phát '
        'hiện chính. Trình bày như EXPLORATORY. Cần '
        'phân tách Brief-COPE 14 nhân tố trong nghiên '
        'cứu tiếp theo.'
    )

    _H(d, '7. Tài liệu tham khảo', 2)
    for r in [
        'Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100. https://doi.org/10.1207/s15327558ijbm0401_6',
        'Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991. https://doi.org/10.1037/bul0000110',
        'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
    ]:
        _ref(d, r)

    out = OUT_DIR / 'Dien_giai_yeu_to_BPDP_bien_phap_doi_pho.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')


# Run all
print('Building 4 docs:')
build_GBTH()
build_HTCM()
build_HTBB()
build_BPDP()
print('Done.')
