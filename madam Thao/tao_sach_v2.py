# -*- coding: utf-8 -*-
"""
Sách mỏng v2: Những mẩu chuyện về Madame Thảo
Mở rộng từ transcript + tham khảo báo chí
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.6
for s in doc.sections:
    s.top_margin = Cm(3)
    s.bottom_margin = Cm(3)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2.5)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def p(text, italic=False, size=13):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.italic = italic

def quote(text):
    para = doc.add_paragraph()
    para.paragraph_format.left_indent = Cm(1)
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.italic = True

def note(text):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(128, 128, 128)

def sep():
    p('\u2015' * 20, size=10)

# ===== BÌA =====
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_heading('', level=0)
r = title.add_run('NHỮNG MẨU CHUYỆN\nVỀ MADAME THẢO')
r.font.name = 'Times New Roman'
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph()
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('TS. Nguyễn Thị Phương Thảo\nChủ tịch HĐQT Tập đoàn Sovico\nChủ tịch HĐQT Vietjet')
r.font.name = 'Times New Roman'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(100, 100, 100)
doc.add_paragraph()
doc.add_paragraph()
epi = doc.add_paragraph()
epi.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = epi.add_run('\u201cMỗi con người đều có thể trở nên vĩ đại.\nLàm những việc lớn có ý nghĩa,\nmà từ những hoạt động rất là đơn giản, giản dị.\u201d')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.italic = True
doc.add_page_break()

# ===== LỜI MỞ ĐẦU =====
h('Lời mở đầu')
p('Quyển sách nhỏ này tập hợp những câu chuyện do chính TS. Nguyễn Thị Phương Thảo \u2014 Chủ tịch HĐQT Tập đoàn Sovico và Vietjet \u2014 kể lại trong một cuộc họp nội bộ về chương trình đào tạo lãnh đạo.')
p('Trong cuộc họp ấy, một thành viên nhóm đã đề xuất: \u201cMadame cứ nói đi, bây giờ mình quay phim. Các bạn lo bản chân, bản mang, phim và câu chuyện.\u201d Bởi vì, như chính Madame Thảo nói, \u201cBây giờ có lẽ là lần đầu tiên mình mới thống kê lại.\u201d')
p('Những câu chuyện ấy \u2014 từ cô sinh viên 17 tuổi một mình sang Nga, đến nữ tỷ phú tự thân đầu tiên của Việt Nam và Đông Nam Á \u2014 không chỉ là chuyện kinh doanh. Đó là câu chuyện về những giá trị mà một con người có thể mang lại cho đất nước mình, cho hàng triệu con người chưa từng gặp mặt.')
quote('\u201cTất cả những hoạt động của mình cho người dân, nó rất là thiết thực, ý nghĩa. Nó xuyên suốt mấy chục năm. Bây giờ mình phải lan tỏa tinh thần ấy cho nhân viên, cán bộ của mình, để mọi người tự hào, mọi người tự tin rằng mình sẽ làm được, bất chấp tất cả những khó khăn, thách thức từ bên ngoài.\u201d')
doc.add_page_break()

# ===== CHƯƠNG 1 =====
h('Cô gái 17 tuổi ở Moscow', 2)
p('Năm 1987, cô gái Hà Nội Nguyễn Thị Phương Thảo, 17 tuổi, một mình lên đường sang Liên Xô du học. Cô theo học ngành Kinh tế Tài chính tại Cao đẳng Kinh tế Quốc dân Moscow và Học viện Thương mại Moscow. Sau này bà hoàn thành luận án Tiến sĩ Điều khiển học Kinh tế.')
p('Ngay từ năm hai đại học, Phương Thảo đã bắt đầu buôn bán \u2014 phân phối máy fax, nhựa và cao su từ Nhật Bản, Hồng Kông và Hàn Quốc sang Liên Xô. Đó là thời kỳ Perestroika, khi nền kinh tế Liên Xô đang chuyển đổi mạnh mẽ và cơ hội kinh doanh mở ra cho những ai dám nắm bắt. Ở tuổi 21, cô đã trở thành triệu phú đô la.')
p('Nhưng câu chuyện không dừng ở việc kiếm tiền. Trong cuộc họp nội bộ, Madame Thảo kể lại hành trình tiếp theo:')
quote('\u201cViệt Nam mình mới thoát ra khỏi tình trạng nợ xấu ấy, bắt đầu đi ra thế giới. Thì mình lại đi tư vấn phát hành những trái phiếu chính phủ đầu tiên ở thị trường chứng khoán New York. Thành công những tỷ đô la đầu tiên. Rồi thành lập những ngân hàng thương mại đầu tiên.\u201d')
p('Bà cùng chồng \u2014 ông Nguyễn Thanh Hùng \u2014 thành lập Sovico Holdings vào cuối những năm 1980. Sau đó, bà góp vốn thành lập Ngân hàng Techcombank (1993) và Ngân hàng VIB (1996) \u2014 những ngân hàng thương mại cổ phần tư nhân đầu tiên của Việt Nam trong thời kỳ đổi mới. Rồi bà bắt đầu xuất khẩu nông sản Việt Nam ra thế giới.')
p('Hãy hình dung: từ một đất nước còn nợ xấu đến trái phiếu chính phủ tại New York, từ máy fax ở Moscow đến ngân hàng thương mại ở Hà Nội \u2014 hành trình ấy diễn ra chỉ trong khoảng mười năm. Và người dẫn dắt hành trình ấy bắt đầu khi mới 17 tuổi.')
note('Nguồn: lời kể từ cuộc họp nội bộ + bazaarvietnam.vn + interstellas.com.vn + mekongasean.vn + dnse.com.vn')
doc.add_page_break()

# ===== CHƯƠNG 2 =====
h('Bao nhiêu tấn thóc để mua một vé máy bay?', 2)
p('Năm 2007, Nguyễn Thị Phương Thảo sáng lập Vietjet Air. Ban đầu, kế hoạch là xây dựng một hãng hàng không 5 sao. Nhưng rồi mọi thứ thay đổi.')
p('Nhưng rồi, trong một chuyến đi lên vùng cao, một cuộc gặp gỡ đã thay đổi tất cả. Một bà mẹ hỏi một câu mà Madame Thảo không bao giờ quên:')
quote('\u201cBao nhiêu tấn thóc để mua một vé máy bay?\u201d')
p('Câu hỏi giản dị ấy đã thay đổi hướng đi của Vietjet. Từ kế hoạch 5 sao, Madame Thảo chuyển hướng hoàn toàn sang mô hình hàng không đại chúng, giá rẻ \u2014 để hàng triệu người Việt Nam lần đầu tiên có thể bay. Mô hình ấy kéo theo sự phát triển mạnh mẽ của kinh tế, du lịch và đầu tư trên khắp cả nước.')
p('Đến nay, Vietjet vận hành hơn 100 máy bay, thực hiện hơn 400 chuyến bay mỗi ngày trên gần 150 đường bay nội địa và quốc tế, phục vụ hơn 200 triệu lượt khách kể từ khi thành lập.')
p('Đến nay, Vietjet đã chiếm hơn 40% thị phần hàng không nội địa. Hàng triệu người Việt Nam \u2014 những người chưa bao giờ nghĩ mình sẽ đi máy bay \u2014 đã bay. Câu hỏi của bà mẹ vùng cao ấy đã được trả lời.')
p('Khi IPO trên sàn chứng khoán tháng 2/2017, Nguyễn Thị Phương Thảo trở thành nữ tỷ phú tự thân đầu tiên của Việt Nam, tỷ phú thứ hai của đất nước (sau Phạm Nhật Vượng), và là nữ tỷ phú duy nhất của Đông Nam Á thời điểm đó. Tạp chí Forbes xếp bà vào danh sách 100 phụ nữ quyền lực nhất thế giới.')
p('Nhưng với Madame Thảo, điều quan trọng hơn danh hiệu là mục đích:')
quote('\u201cTiền nhiều để hiện thực hóa những ước mơ cao đẹp và giúp đỡ được nhiều người hơn.\u201d')
note('Nguồn: tuoitre.vn (9/1/2020) + crystalbay.com + Forbes + Wikipedia')
doc.add_page_break()

# ===== CHƯƠNG 3 =====
h('Nâng đôi cánh cho Tân Sơn Nhất', 2)
p('Sân bay Tân Sơn Nhất \u2014 cửa ngõ hàng không lớn nhất phía Nam \u2014 từng là nỗi ám ảnh của cả ngành hàng không Việt Nam. Quá tải triền miên, chậm chuyến liên tục, hành khách mệt mỏi chờ đợi. Nhiều người nói rằng không thể cải thiện được nữa. Nhưng Madame Thảo không nghĩ vậy:')
quote('\u201cNăm 2013 mình làm, 2014 bắt đầu chạy. Dự án nâng cao năng lực điều hành cất hạ cánh ở sân bay Tân Sơn Nhất \u2014 nâng lên gấp đôi từ 30 lượt đến 60 lượt. Bây giờ đang vận hành 48 lượt. Từ 24 triệu lượt khách lên 42 triệu lượt khách.\u201d')
p('Điều đáng nói là dự án không chỉ phục vụ Vietjet. Toàn bộ ngành hàng không \u2014 tất cả các hãng, tất cả hành khách \u2014 đều được hưởng lợi. ACV (Tổng công ty Cảng hàng không Việt Nam) trong mười mấy năm liền, mỗi năm đón thêm 16 triệu lượt khách. Madame Thảo nói với một chút hài hước:')
quote('\u201cĐồng nghĩa là họ có lợi nhuận mấy trăm triệu đô la một năm. Từ lúc đến giờ mình đã mang lại cho họ bao nhiêu tỷ đô la lợi nhuận. Chứ không phải là doanh số đâu. Thì làm sao mà chẳng có tiền đi xây sân bay.\u201d')
p('Đến năm 2025, nhà ga T3 mới với công suất 20 triệu khách mỗi năm đã nâng tổng công suất Tân Sơn Nhất lên 50 triệu khách mỗi năm. Sân bay vận hành tới 48 lượt cất hạ cánh mỗi giờ vào cao điểm \u2014 nghĩa là cứ 1 phút 15 giây lại có một máy bay cất hoặc hạ cánh.')
note('Nguồn: lời kể từ cuộc họp nội bộ + plo.vn + dantri.com.vn (2025)')
doc.add_page_break()

# ===== CHƯƠNG 4 =====
h('100 ngày giải cứu sàn chứng khoán', 2)
p('Đầu năm 2021, giữa lúc cả nước đang vật lộn với COVID-19, sàn HoSE (Sở Giao dịch Chứng khoán TP.HCM) rơi vào khủng hoảng. Hệ thống giao dịch do Hàn Quốc cung cấp không đáp ứng được khối lượng giao dịch bùng nổ \u2014 nghẽn lệnh triền miên, nhà đầu tư không thể mua bán, hàng nghìn tỷ đồng bị \u201cđóng băng\u201d mỗi ngày. Thị trường vốn của cả nước bị đe dọa.')
p('Trong khi nhiều người còn đang bàn thảo giải pháp, Madame Thảo đã hành động. Bà kết nối Sovico với FPT \u2014 tập đoàn công nghệ hàng đầu Việt Nam \u2014 để xây dựng hệ thống giao dịch thay thế. Trong cuộc họp nội bộ, bà kể lại tầm vóc của sự việc bằng một so sánh khiến người nghe phải giật mình:')
quote('\u201cCả nước Mỹ có sàn New York, Nasdaq \u2014 2 cái sàn chứng khoán. Cả nước Anh có London Stock Exchange. Mà nó bị treo thì giải cứu mấy trăm tỷ đô la ở đấy. Có hơn 100 công ty chứng khoán, 4 triệu khách hàng. Mỗi ngày giao dịch khoảng 1 tỷ đô la, mà cao là 2 tỷ đô la. Mình cũng nhảy vào giải cứu.\u201d')
p('Madame Thảo cho biết bà đã liên hệ trực tiếp với lãnh đạo các tập đoàn công nghệ lớn như FPT, và nhận được phản hồi rằng chỉ cần khoảng 2 tháng với chi phí khoảng 60 tỷ đồng là có thể giải quyết. Một con số khiêm tốn so với hàng trăm tỷ đô la đang bị đe dọa mỗi ngày trên sàn.')
p('Chỉ trong khoảng 100 ngày, hệ thống giao dịch mới do FPT xây dựng chính thức vận hành ngày 5 tháng 7 năm 2021 \u2014 giải cứu thành công thị trường vốn cho cả nước. Thanh khoản mở lại, dòng tiền chảy vào thị trường.')
p('Kết quả: Chính phủ giữ được sàn. Hơn 100 công ty chứng khoán hoạt động trở lại. 4 triệu khách hàng được giao dịch bình thường. Các công ty niêm yết có thị trường vốn để huy động và phát triển. Thanh khoản mở lại, dòng tiền chảy vào thị trường.')
p('Và phần thưởng cho Sovico?')
quote('\u201cThế còn mình thì không cần cái gì cả.\u201d')
note('Nguồn: lời kể từ cuộc họp nội bộ + tinnhanhchungkhoan.vn + tuoitre.vn (5/7/2021) + cafef.vn')
doc.add_page_break()

# ===== CHƯƠNG 5 =====
h('Ngoại giao trên bầu trời', 2)
p('Mỗi hợp đồng mua máy bay trị giá hàng tỷ đô la không chỉ là giao dịch thương mại \u2014 đó là công cụ ngoại giao kinh tế. Madame Thảo giải thích:')
quote('\u201cMua máy bay Boeing là đi với nước Mỹ \u2014 ông Trump, ông Obama, ông Bush. Mua máy bay Airbus là đi với Pháp, Đức, Tây Ban Nha \u2014 là những nước sản xuất tàu bay. Mua máy bay COMAC của Trung Quốc là chương trình của ông Tập Cận Bình. Là mình tạo nên xây dựng mối quan hệ ngoại giao, xây dựng kinh tế giữa các quốc gia, dân tộc. Mình sẽ giảm bớt đi những chiến tranh, những sự xung đột.\u201d')
p('Những con số cho thấy tầm vóc: Tháng 2/2019, trong chuyến thăm Hà Nội của Tổng thống Trump, Vietjet ký hợp đồng 100 chiếc Boeing 737 MAX trị giá 12,7 tỷ USD \u2014 một trong những đơn hàng máy bay lớn nhất châu Á thời điểm đó. Lễ ký kết diễn ra với sự chứng kiến của cả hai nguyên thủ. Tháng 2/2026, trong chuyến thăm Washington của Tổng Bí thư Tô Lâm, Vietjet tiếp tục ký thêm các thỏa thuận mua Boeing. Mỗi hợp đồng không chỉ mang về máy bay \u2014 mà mang về thiện chí giữa hai quốc gia.')
p('Và không chỉ là hợp đồng:')
quote('\u201cMáy bay mình bay đến đâu, biên giới đất nước nó mở ra. Mình mang cờ cho khắp thế giới, các đường bay mình đến.\u201d')
sep()
p('Ngày 15 tháng 4 năm 2021, Đại sứ Pháp tại Việt Nam Nicolas Warnery đã trao cho Madame Thảo Huân chương Bắc đẩu Bội tinh (Ordre national de la Légion d\'honneur) \u2014 huân chương cao quý nhất của nước Pháp, được Napoléon Bonaparte thiết lập năm 1802. Bà là nữ doanh nhân Việt Nam đầu tiên nhận vinh dự này, vì những đóng góp trong thúc đẩy quan hệ Việt-Pháp.')
p('Bà cũng được Đại học Harvard đưa vào tài liệu nghiên cứu và giảng dạy, được Business Insider Australia bầu chọn là một trong 100 người thay đổi kinh tế châu Á, và được tạp chí Tatler vinh danh là người có ảnh hưởng nhất châu Á trong lĩnh vực từ thiện.')
note('Nguồn: lời kể từ cuộc họp nội bộ + Bloomberg (27/2/2019) + CNBC (19/2/2026) + baochinhphu.vn + vnexpress.net (15/4/2021)')
doc.add_page_break()

# ===== CHƯƠNG 6 =====
h('90 ngày \u2014 khi Hàn Quốc mất 17 năm', 2)
p('Năm 2017, Việt Nam đăng cai APEC \u2014 Diễn đàn Hợp tác Kinh tế Châu Á-Thái Bình Dương \u2014 tại Đà Nẵng. Đây là sự kiện ngoại giao lớn nhất trong lịch sử Việt Nam thời điểm đó, với sự tham dự của Tổng thống Trump, Chủ tịch Tập Cận Bình, Tổng thống Putin, Thủ tướng Shinzo Abe và lãnh đạo 21 nền kinh tế.')
p('Sovico và Vietjet đã tham gia phục vụ sự kiện. Madame Thảo kể lại một so sánh ấn tượng về năng lực triển khai:')
quote('\u201cCái giải pháp chiếu sáng mình làm có 90 ngày thôi. Bình thường nó phải mất 12 tháng mới xong. Trong khi đó Hàn Quốc triển khai 12 năm mà không chạy, đến lúc bàn giao đã 17 năm. Tốn 100 triệu đô rồi, không bàn giao được.\u201d')
quote('\u201cXây cung hội nghị 4 tháng.\u201d')
quote('\u201cNhư là cái tòa Marina này, nó đã mang đến một cái sinh khí cho cả cái thành phố này.\u201d')
p('Bà cũng kể về một chuyến đi đặc biệt:')
quote('\u201cNgày 9 tháng 1 năm ngoái, không ai mời mình đâu, nhưng mà thấy thời cuộc nó căng thẳng. Mình có mối quan hệ tốt với ông Trump, Đại sứ quán Mỹ mời mình, thì mình bay luôn, mang máy bay đi luôn. Cũng tạo nên thiện cảm, cũng đỡ căng thẳng.\u201d')
p('Đóng góp cho APEC được Madame Thảo mô tả bằng tám chữ: \u201cĐưa Việt Nam vào thế giới và đưa thế giới về Việt Nam.\u201d')
note('Nguồn: lời kể từ cuộc họp nội bộ. APEC 2017 tại Đà Nẵng xác nhận qua vnexpress.net + baochinhphu.vn. Chi tiết \u201c90 ngày\u201d và \u201cHàn Quốc 17 năm\u201d là từ lời kể, cần xác nhận thêm.')
doc.add_page_break()

# ===== CHƯƠNG 7 =====
h('Mùa COVID \u2014 lao đi vì người dân', 2)
p('Giữa năm 2021, đại dịch COVID-19 tấn công TP.HCM khốc liệt. Thành phố phong tỏa hoàn toàn. Bệnh viện quá tải. Số ca mắc mới lên đến hơn 10.000 ca mỗi ngày vào tháng 9. Hàng triệu người thiếu lương thực.')
p('Trong bối cảnh ấy, Madame Thảo không ngồi chờ. Bà kể, giọng nhẹ nhàng nhưng dứt khoát:')
quote('\u201cMùa Covid mình lao đi mình chở vắc-xin, rồi thiết bị y tế, bao nhiêu những bao gạo. Không có những cái đó thì làm sao mà người ta được yên lòng.\u201d')
p('Theo báo chí ghi nhận, Vietjet đã chuyên chở miễn phí hơn 2.000 bác sĩ và nhân viên y tế từ miền Bắc và miền Trung vào TP.HCM. Hàng trăm máy thở và thiết bị y tế từ Đức được vận chuyển miễn phí về Việt Nam.')
p('Madame Thảo khởi xướng chương trình \u201cBữa cơm yêu thương\u201d \u2014 mỗi ngày cung cấp 20.000 suất cơm chay và mặn miễn phí cho bác sĩ tại bệnh viện dã chiến, khu cách ly, và các gia đình khó khăn. Trong hơn một tháng, chương trình đã cung cấp hơn 600.000 suất ăn và hàng chục tấn thực phẩm.')
p('Vietjet và HDBank tặng 10 xe cứu thương cho Sở Y tế TP.HCM, tài trợ trang thiết bị cho hơn 1.000 giường hồi sức tích cực. Tổng cộng ước tính hơn 70 tỷ đồng cho công tác thiện nguyện và phòng chống dịch.')
p('70 tỷ đồng. 600.000 suất ăn. 2.000 bác sĩ được chở miễn phí. 10 xe cứu thương. Hơn 1.000 giường hồi sức. Những con số ấy, nếu chỉ đọc trên báo, có thể chỉ là thống kê. Nhưng với những bác sĩ trong bệnh viện dã chiến, với những gia đình không có gạo ăn giữa mùa phong tỏa \u2014 mỗi suất cơm, mỗi bao gạo đều là điều họ cần nhất lúc đó.')
p('Báo Tuổi Trẻ đã đặt tiêu đề: \u201cNữ tỷ phú hỗ trợ nóng hơn 70 tỷ đồng cho những ngày giãn cách ở TP.HCM.\u201d')
note('Nguồn: lời kể + vietnamplus.vn + thanhnien.vn + tuoitre.vn (22/7/2021) + hdbank.com.vn')
doc.add_page_break()

# ===== CHƯƠNG 8 =====
h('Cái thang máy dành cho quan tài', 2)
p('Trong cuộc họp, Madame Thảo chia sẻ một triết lý thiết kế mà ít ai biết:')
quote('\u201cTất cả những cái gì mà mình làm \u2014 thiết kế một cái nhà chung cư, là phải có cái thang máy lớn để dành cho cáng cứu thương và quan tài. Từ cái thiết kế là mình cũng nghĩ đến con người \u2014 lúc người ta bệnh, người ta cần phải đưa đi cáng gì chứ. Một cái thang máy ở nhà chung cư mà không thiết kế ra như thế.\u201d')
quote('\u201cỞ đâu có bậc thang là mình phải có cái đường dốc cho xe nôi, cho xe người khuyết tật. Đấy là tinh thần mình đưa vào mọi hoạt động.\u201d')
p('Hãy dừng lại một giây để nghĩ về điều này. Khi thiết kế chung cư, hầu hết kiến trúc sư không nghĩ đến chuyện đưa quan tài xuống thang máy. Không nghĩ đến cáng cứu thương cần bao nhiêu không gian. Không nghĩ đến xe nôi và xe lăn cần đường dốc ở mỗi bậc thang. Madame Thảo nghĩ đến tất cả những điều ấy.')
p('Đó là cách bà nghĩ về con người \u2014 không phải ở những lúc vui vẻ, mà ở những lúc con người cần được giúp đỡ nhất.')
p('Phú Long \u2014 công ty bất động sản thuộc hệ sinh thái Sovico \u2014 đã áp dụng tinh thần này vào các dự án nhà ở. Tinh thần ấy cũng xuất hiện trong thiết kế máy bay Vietjet, trong bố trí trường học Victoria, trong kiến trúc các tổ hợp du lịch nghỉ dưỡng.')
note('Nguồn: lời kể từ cuộc họp nội bộ.')
doc.add_page_break()

# ===== CHƯƠNG 9 =====
h('10.000 đôi mắt mỗi năm', 2)
p('Khi được hỏi về thành tựu xã hội của Sovico và Vietjet, Madame Thảo kể ra một danh sách dài đến mức chính bà cũng thừa nhận: \u201cBây giờ có lẽ là lần đầu tiên mình mới thống kê lại.\u201d')
p('Một thành viên trong cuộc họp bổ sung:')
quote('\u201cRiêng cái chương trình mổ mắt cho người nghèo \u2014 một năm 10.000 ca mổ mắt. Rồi tổ chức khám chữa bệnh ở khắp nơi.\u201d')
p('Madame Thảo tiếp tục liệt kê:')
quote('\u201cCờ vua Việt Nam lên top 7 thế giới. Làm bóng đá futsal. Rồi bóng đá nữ, rồi thể thao người khuyết tật.\u201d')
p('Giải cờ vua quốc tế HDBank \u2014 do chính Madame Thảo khởi xướng gần một thập niên trước \u2014 đã trở thành một trong những giải đấu cờ vua uy tín nhất châu Á. Năm 2021, Liên đoàn Cờ Thế giới (FIDE) ký thỏa thuận hợp tác 10 năm với HDBank (2022\u20132031). Nhờ nền tảng ấy, cờ vua Việt Nam đã sản sinh ra Đại kiện tướng Lê Quang Liêm \u2014 từng lọt top 20 thế giới và đoạt cúp vô địch cờ nhanh thế giới.')
p('Rồi những phong trào mới lạ:')
quote('\u201cKêu gọi cho bà con bắt đầu đưa hàng hóa lên trên TikTok để bán. Thực ra mình là những cái người đi phát động.\u201d')
p('Dự án \u201cTôi yêu Tổ quốc tôi\u201d với hơn 6 triệu đoàn viên thanh niên tham gia. Các chương trình thiện nguyện tại Philippines, Ấn Độ. Vietjet sử dụng máy bay mới để giảm khí thải, chống biến đổi khí hậu. HDBank chuyển đổi sang ngân hàng xanh. Sovico hợp tác với Đại học Oxford (trị giá 155 triệu bảng Anh) để phát triển nghiên cứu và giáo dục.')
note('Nguồn: lời kể + dantri.com.vn + vietnamchess.vn + nhandan.vn + tuoitre.vn + baochinhphu.vn')
doc.add_page_break()

# ===== CHƯƠNG 10 =====
h('Cái hồn của đào tạo', 2)
p('Nhưng giữa tất cả những thành tựu lớn lao ấy \u2014 tỷ đô, giải cứu, ngoại giao, huân chương \u2014 Madame Thảo lại trăn trở nhất về một điều rất gần, rất nhỏ: tinh thần của chính những người đang đi cùng mình mỗi ngày.')
quote('\u201cĐấy là cái tinh thần mà mình đưa vào cho cái đào tạo của mình. Nó khác với những cái hoạt động thông thường mà sau này mình cứ làm cho nó mai một đi, làm cho nó đi đâu mất. Tức là nó thiếu cái kiểu giống như con người, nó thiếu cái hồn của nó.\u201d')
p('Bà kể lại hình ảnh tổ bay ngày xưa:')
quote('\u201cNgày xưa tổ bay là briefing ở tại trụ sở công ty. Thậm chí là còn xếp hàng ngay ngắn. Đi một lượt, kiểm tra lại với nhau, xong rồi diễu ở trong văn phòng ấy. Nhân viên đầu tiên nhìn cái tổ bay này chuẩn bị làm nhiệm vụ, thì người ta tự hào, người ta thấy cái giá trị, cái ý nghĩa của công việc.\u201d')
p('Rồi bà quay sang thực tế hiện tại:')
quote('\u201cChứ còn bây giờ là cứ đầu tắt mặt tối, ra khỏi nhà 2\u20133 giờ sáng, ra sân bay rồi về đi ngủ. Kiếm một số tiền thì không nhỏ, nhưng mà cũng đủ sống. Không có thời gian tiêu. Người ta rời đi không phải vì chế độ của mình. Chế độ của mình là chế độ tốt. Nhưng mà cái môi trường, người ta cảm thấy là không có ý nghĩa.\u201d')
p('Chính vì điều đó, Madame Thảo quyết định xây dựng lại chương trình đào tạo lãnh đạo \u201cHội nhập Tỏa sáng\u201d \u2014 không phải để dạy kỹ năng, mà để lấy lại cái hồn:')
quote('\u201cTừ lãnh đạo thì mới lan tỏa xuống nhất nhất. Tinh thần ấy mới xuống.\u201d')
note('Nguồn: lời kể từ cuộc họp nội bộ.')
doc.add_page_break()

# ===== CHƯƠNG 11 =====
h('Nhân quả', 2)
p('Triết lý quản lý của Madame Thảo không đến từ sách vở hay trường kinh doanh:')
quote('\u201cMình tin vào nhân quả. Nhân nào quả đấy. Mình làm chưa tốt thì cuộc sống sẽ có cái quả tương ứng. Chứ không cần phải có hành động gì xử lý căng thẳng. Xử lý ở trong nội bộ thôi.\u201d')
p('Về quyền được sai:')
quote('\u201cCuộc sống, người ta là con người, người ta có quyền sai. Thì mình hãy tạo cho người ta cơ hội để có thể khắc phục.\u201d')
p('Về công bằng:')
quote('\u201cTạo nên sự công bằng trong tổ chức là cần thời gian, cần cơ chế. Chứ không phải bây giờ cấp dưới đụng chạm một cái gì đấy mà mình cho rằng là sai, thì mình xử lý rất là mạnh. Trong khi đó lãnh đạo ở trên không gương mẫu, thì đấy nó không công bằng.\u201d')
p('Về vai trò lãnh đạo:')
quote('\u201cMỗi người phải trở thành người coach cho những người xung quanh. Cũng phải biết cách học hỏi.\u201d')
p('Và về kỹ năng còn thiếu:')
quote('\u201cKỹ năng mà tổ chức mình đang thiếu: lập kế hoạch, quản lý dự án \u2014 là những cái cơ bản. Mình không biết được những cái gì cần thiết vì mình chưa được đào tạo. Thì đào tạo xong mình sẽ thiết kế được thôi. Chả có gì mà khó.\u201d')
p('Đó là Madame Thảo \u2014 người phụ nữ tin vào nhân quả, tin rằng mọi người đều có quyền sai, tin rằng công bằng cần hệ thống chứ không phải cảm tính, và tin rằng không có gì là khó nếu được đào tạo bài bản. Giản dị đến mức khiến người ta phải suy nghĩ lại về những gì mình vẫn cho là phức tạp.')
note('Nguồn: lời kể từ cuộc họp nội bộ + tuoitre.vn (9/1/2020)')
doc.add_page_break()

# ===== KẾT =====
h('Lời kết')
p('Những câu chuyện trong cuốn sách nhỏ này chỉ là một phần rất nhỏ trong hành trình mấy chục năm của Madame Thảo.')
p('Từ cô sinh viên ở Moscow đến nữ tỷ phú Forbes. Từ trái phiếu chính phủ tại New York đến giải cứu sàn chứng khoán TP.HCM. Từ ngoại giao trên bầu trời với Trump, Tập Cận Bình đến 10.000 ca mổ mắt cho người nghèo mỗi năm. Từ \u201cBao nhiêu tấn thóc để mua vé bay?\u201d đến hãng hàng không chở hơn 200 triệu lượt khách. Từ cái thang máy dành cho quan tài đến 600.000 suất cơm yêu thương giữa mùa dịch.')
p('Tất cả đều hướng về một thông điệp:')
quote('\u201cMỗi con người đều có thể trở nên vĩ đại. Làm những việc lớn có ý nghĩa, mà từ những hoạt động rất là đơn giản, giản dị của mình thôi.\u201d')
doc.add_paragraph()
sep()
p('Sách được biên soạn từ bản ghi cuộc họp nội bộ (tháng 3/2026), có tham khảo thêm nguồn báo chí. Các trích dẫn trực tiếp trong ngoặc kép là lời kể nguyên văn của Madame Thảo, giữ nguyên giọng nói tự nhiên. Một số chi tiết chỉ có trong lời kể (đánh dấu trong ghi chú từng chương) cần được xác nhận từ nguồn chính thức trước khi sử dụng trong tài liệu đối ngoại.', size=10, italic=True)

fname = 'Sách - Những mẩu chuyện về Madame Thảo.docx'
doc.save(fname)
print('Done: ' + fname)
