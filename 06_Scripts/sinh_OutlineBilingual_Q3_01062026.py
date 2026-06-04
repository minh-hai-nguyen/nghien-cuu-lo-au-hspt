# -*- coding: utf-8 -*-
"""Sinh outline song ngu Q3 PLOS ONE de gui co-authors.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'OutlineBilingual_Q3_01062026.docx')

d = Document()
for sec in d.sections:
    sec.page_width = Cm(21.0); sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.25


def H1(text_vn, text_en=None):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text_vn); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    if text_en:
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run(text_en); r2.font.name = 'Times New Roman'; r2.font.size = Pt(13); r2.italic = True
        r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H2(text_vn, text_en=None):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text_vn); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    if text_en:
        r2 = p.add_run('  /  ' + text_en); r2.font.name = 'Times New Roman'
        r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

def H3(text_vn, text_en=None):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text_vn); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    if text_en:
        r2 = p.add_run('  /  ' + text_en); r2.font.name = 'Times New Roman'
        r2.font.size = Pt(10); r2.italic = True; r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

def Para(text_vn, text_en=None, indent=True, italic=False):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text_vn); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic
    if text_en:
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.space_after = Pt(8)
        if indent: p2.paragraph_format.first_line_indent = Cm(0.5)
        r2 = p2.add_run(text_en); r2.font.name = 'Times New Roman'; r2.font.size = Pt(10); r2.italic = True
        r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def Bullet(text_vn, text_en=None, level=0):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('▸ ' + text_vn); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    if text_en:
        p2 = d.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.left_indent = Cm(0.9 + level * 0.5); p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(text_en); r2.font.name = 'Times New Roman'; r2.font.size = Pt(10); r2.italic = True
        r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def NCS_marker(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ [NCS/Thầy confirm] ' + text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True; r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_table(headers_vn, headers_en, rows, col_widths_cm, style='Light Grid Accent 1'):
    t = d.add_table(rows=1, cols=len(headers_vn))
    t.style = style; t.autofit = False
    hdr = t.rows[0].cells
    for i, h_vn in enumerate(headers_vn):
        h_en = headers_en[i] if headers_en and i < len(headers_en) else None
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h_vn); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        if h_en:
            p.add_run('\n')
            r2 = p.add_run(h_en); r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.italic = True
            r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            if isinstance(cell_data, tuple):
                vn, en = cell_data
                r = p.add_run(vn); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if en:
                    p.add_run('\n')
                    r2 = p.add_run(en); r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.italic = True
                    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            else:
                r = p.add_run(str(cell_data)); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
# COVER
# ============================================================
H1('ĐỀ CƯƠNG CHI TIẾT BÀI BÁO Q3', 'DETAILED OUTLINE — Q3 PAPER')

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Bản song ngữ Tiếng Việt – English  /  Bilingual Edition')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

d.add_paragraph()
Para('Tài liệu này là đề cương chi tiết bài báo nghiên cứu mô tả gửi tới '
     'tạp chí Q1/Q3 quốc tế PLOS ONE. Bản outline được trình bày song ngữ '
     'Tiếng Việt – English để thuận tiện cho việc rà soát của đồng tác giả.',
     'This document presents the detailed outline of a descriptive research '
     'paper to be submitted to PLOS ONE. The bilingual Vietnamese–English '
     'format is designed for review by co-authors.', indent=False)


key_info = [
    (('Tên bài (tentative)', 'Title (tentative)'),
     ('Biểu hiện và mô hình các dạng rối loạn lo âu ở học sinh trung học '
      'cơ sở Việt Nam: Nghiên cứu mô tả cắt ngang',
      'Manifestations and patterns of anxiety disorder subtypes among '
      'Vietnamese lower secondary school students: A descriptive cross-'
      'sectional study')),
    (('Tạp chí mục tiêu', 'Target journal'),
     ('PLOS ONE (Q1/Q2 by SJR, IF 3.7, acceptance ~50%)',
      'PLOS ONE (Q1/Q2 by SJR, IF 3.7, ~50% acceptance rate)')),
    (('Backup journal', 'Backup journal'),
     ('BMC Pediatrics (IF 2.5) nếu cần pediatric focus',
      'BMC Pediatrics (IF 2.5) if pediatric focus is needed')),
    (('Số từ dự kiến', 'Target word count'),
     ('3.500–5.000 từ', '3,500–5,000 words')),
    (('Tác giả', 'Authors'),
     ('Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹†',
      'Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹†')),
    (('Ghi chú tác giả', 'Author notes'),
     ('*: Tác giả thứ nhất, †: Tác giả liên hệ',
      '*: First author, †: Corresponding author')),
]
make_table(['Mục', 'Nội dung'], ['Item', 'Content'], key_info, [4.0, 13.0])

d.add_page_break()


# ============================================================
H1('1. TÓM TẮT KHOA HỌC', '1. SCIENTIFIC ABSTRACT')

Para('Abstract chuẩn cấu trúc PLOS ONE (~200 từ).',
     'Structured abstract per PLOS ONE guidelines (~200 words).',
     indent=False, italic=True)

H3('Background', 'Background')
Para('Rối loạn lo âu ở học sinh THCS Việt Nam còn thiếu dữ liệu chuẩn '
     'mức độ mục (item-level normative data) cần thiết để giải thích lâm '
     'sàng và phát triển công cụ sàng lọc. Việc phân tích chi tiết các '
     'dạng RLLA theo DSM-5 cung cấp cơ sở cho can thiệp đặc thù.',
     'Vietnamese lower secondary school students lack item-level normative '
     'anxiety data essential for clinical interpretation and screening tool '
     'development. Detailed analysis by DSM-5 anxiety disorder subtypes '
     'provides basis for targeted intervention.')

H3('Methods', 'Methods')
Para('Nghiên cứu mô tả cắt ngang trên 1.352 học sinh THCS Hà Nội (614 nam, '
     '738 nữ; khối lớp 6-9). Phiên bản tiếng Việt của RCADS (Chorpita, 2000) '
     'gồm 15 mục đo 3 dạng RLLA (GAD 7 mục, SAD 4 mục, SocAD 4 mục). '
     'Phân tích bao gồm thống kê mô tả, kiểm định t-test/ANOVA với hiệu '
     'chỉnh Bonferroni, và cỡ tác động (Cohen d, η²).',
     'Cross-sectional descriptive study on 1,352 lower secondary school '
     'students in Hanoi (614 males, 738 females; grades 6-9). Vietnamese '
     'RCADS adaptation (Chorpita, 2000) with 15 items measuring 3 anxiety '
     'subtypes (GAD 7, SAD 4, SocAD 4). Analyses included descriptive '
     'statistics, t-test/ANOVA with Bonferroni correction, and effect sizes '
     '(Cohen d, η²).')

H3('Results', 'Results')
Para('Mức trung bình rối loạn lo âu: GAD M = 55,82; SAD M = 25,06; '
     'SocAD M = 48,41 (thang 0-100). Mục có điểm cao nhất: RCADS4 '
     '(M = 64,28) cho GAD; RCADS32 (M = 56,98) cho SocAD; RCADS46 '
     '(M = 27,88) cho SAD. Khác biệt giới có ý nghĩa cho GAD và SocAD '
     '(F > 44; p < 0,001) nhưng không cho SAD (F = 0,246; p = 0,620). '
     'Quỹ đạo theo khối lớp: SAD giảm đơn điệu từ M = 32,13 (khối 6) → '
     '19,46 (khối 9); GAD tăng nhẹ.',
     'Mean anxiety levels: GAD M = 55.82; SAD M = 25.06; SocAD M = 48.41 '
     '(0-100 scale). Highest-scoring items: RCADS4 (M = 64.28) for GAD; '
     'RCADS32 (M = 56.98) for SocAD; RCADS46 (M = 27.88) for SAD. Gender '
     'differences significant for GAD and SocAD (F > 44; p < 0.001) but not '
     'for SAD (F = 0.246; p = 0.620). Grade trajectory: SAD declines '
     'monotonically from M = 32.13 (grade 6) → 19.46 (grade 9); GAD '
     'increases slightly.')

H3('Conclusions', 'Conclusions')
Para('Nghiên cứu cung cấp dữ liệu chuẩn mức độ mục đầu tiên cho phiên bản '
     'tiếng Việt của RCADS trên mẫu lớn HS THCS. Các mục có điểm cao nhất '
     'được đề xuất làm mục tiêu cho công cụ sàng lọc ngắn. Quỹ đạo phát '
     'triển khác nhau giữa các dạng RLLA gợi ý cần can thiệp theo độ tuổi.',
     'This study provides first item-level normative data for the Vietnamese '
     'RCADS adaptation in a large lower secondary school sample. Highest-'
     'scoring items are proposed as targets for brief screening tools. '
     'Differential developmental trajectories across subtypes suggest the '
     'need for age-appropriate intervention.')


d.add_page_break()


# ============================================================
H1('2. ĐIỂM MỚI CỦA NGHIÊN CỨU', '2. NOVEL CONTRIBUTIONS')

H2('Ba đóng góp chính', 'Three Key Contributions')

contrib_data = [
    (('1', '1'),
     ('Phân tích mức độ mục đầu tiên cho RCADS Việt Nam',
      'First item-level analysis for Vietnamese RCADS adaptation'),
     ('Nghiên cứu đầu tiên cung cấp phân tích chi tiết mức độ mục '
      '(item-level) cho phiên bản tiếng Việt rút gọn của RCADS (Chorpita, '
      '2000) trên cỡ mẫu lớn N = 1.352 học sinh THCS, cung cấp DỮ LIỆU CHUẨN '
      '(normative data) cho lâm sàng tâm lý Việt Nam.',
      'First study to provide detailed item-level analysis for the Vietnamese '
      'adapted RCADS (Chorpita, 2000) on a large N = 1,352 lower secondary '
      'school sample, providing NORMATIVE DATA for Vietnamese clinical '
      'psychology.')),
    (('2', '2'),
     ('Quỹ đạo phát triển cắt ngang theo khối lớp',
      'Cross-sectional developmental trajectory by grade'),
     ('Mô tả mẫu hình thay đổi điểm trung bình rối loạn lo âu từ khối 6 '
      'đến khối 9. Phát hiện rối loạn lo âu chia ly GIẢM đơn điệu '
      '(32,13 → 27,14 → 20,88 → 19,46), phù hợp với phân loại DSM-5 SAD = '
      'khởi phát thời thơ ấu. RLLA lan toả tăng nhẹ qua các khối.',
      'Describes the changing pattern of anxiety mean scores from grade 6 to '
      'grade 9. Finding that separation anxiety DECLINES monotonically '
      '(32.13 → 27.14 → 20.88 → 19.46), consistent with DSM-5 SAD = '
      'childhood-onset classification. Generalized anxiety slightly '
      'increases across grades.')),
    (('3', '3'),
     ('Xác định mục có khả năng dùng làm công cụ sàng lọc',
      'Identification of items as screening targets'),
     ('Các mục có điểm trung bình cao nhất theo từng dạng RLLA được đề '
      'xuất làm cơ sở cho công cụ sàng lọc ngắn (3-5 mục): RCADS4 + RCADS13 '
      '+ RCADS8 (sàng lọc GAD); RCADS32 + RCADS43 (sàng lọc SocAD). '
      'Đây là cơ sở khoa học cho phát triển công cụ sàng lọc trường học '
      'tại Việt Nam.',
      'Items with highest mean scores per subtype are proposed as basis for '
      'brief screening tools (3-5 items): RCADS4 + RCADS13 + RCADS8 (GAD '
      'screener); RCADS32 + RCADS43 (SocAD screener). This provides '
      'scientific basis for developing school-based screening tools in '
      'Vietnam.')),
]
make_table(['#', 'Tiêu đề', 'Mô tả chi tiết'],
           ['#', 'Title', 'Detailed description'],
           contrib_data, [1.0, 4.5, 11.5])


d.add_page_break()


# ============================================================
H1('3. PHẦN GIỚI THIỆU', '3. INTRODUCTION')


H2('§1. Gánh nặng RLLA + Tầm quan trọng phân loại DSM-5',
   '§1. Anxiety burden + DSM-5 subtype importance')

Para('Rối loạn lo âu ảnh hưởng đến 7-15% thanh thiếu niên toàn cầu '
     '(Anderson và cs., 2025); WHO xếp lo âu và trầm cảm vào top nguyên '
     'nhân khuyết tật ở lứa tuổi 10-19. Báo cáo GBD ASEAN 2025 ghi nhận '
     'tỷ lệ 10,1% tại Việt Nam và 16,3% gánh nặng DALYs ở nhóm tuổi 10-14.',
     'Anxiety disorders affect 7-15% of adolescents globally '
     '(Anderson et al., 2025); WHO ranks anxiety and depression among top '
     'disability causes in the 10-19 age group. GBD ASEAN 2025 reports '
     '10.1% prevalence in Vietnam and 16,3% DALY burden in the 10-14 age '
     'group.')

Para('Tại sao việc phân tích theo DSM-5 SUBTYPES quan trọng: GAD, SAD và '
     'SocAD có các mẫu hình khởi phát phát triển khác nhau (Chorpita, '
     '2000), đáp ứng điều trị khác nhau (tập trung nhận thức vs. hành vi), '
     'và chiến lược phòng ngừa khác nhau. Coi "lo âu" như đơn nhất sẽ làm '
     'mờ đi những phân biệt quan trọng này.',
     'Why DSM-5 SUBTYPES matter: GAD, SAD, and SocAD have different '
     'developmental onset patterns (Chorpita, 2000), treatment responses '
     '(cognitive vs. behavioral focus), and prevention strategies. Treating '
     '"anxiety" as monolithic obscures these critical distinctions.')

Para('Lý do phân tích mức độ mục: dữ liệu mức độ mục xác định MỤC TIÊU '
     'TRIỆU CHỨNG cụ thể để sàng lọc và thiết kế can thiệp nhận thức – '
     'hành vi (Chorpita 2000 — RCADS validation).',
     'Item-level analysis rationale: Item-level data identifies SPECIFIC '
     'symptom targets for screening and cognitive-behavioral intervention '
     'design (Chorpita 2000 — RCADS validation).')


H2('§2. Khoảng trống nghiên cứu Việt Nam và bối cảnh V-NAMHS',
   '§2. Vietnamese research gap and V-NAMHS context')

Para('Nghịch lý V-NAMHS 2022: Điều tra quốc gia báo cáo tỷ lệ rối loạn '
     'lo âu chỉ 2,3% sử dụng DISC-5 — một con số đáng chú ý thấp. Trong '
     'khi đó, Hoàng Trung Học và cộng sự (2025) với N = 8.473 sau '
     'COVID-19 báo cáo tỷ lệ cao hơn nhiều khi sử dụng các thang đo DASS. '
     'Sự khác biệt phản ánh lựa chọn công cụ đo lường.',
     'V-NAMHS 2022 paradox: The national survey reports anxiety disorder '
     'prevalence at only 2.3% using DISC-5 — a remarkably low figure. '
     'Meanwhile, Hoang Trung Hoc et al. (2025) with N = 8,473 post-COVID-19 '
     'reports much higher prevalence using DASS measures. The discrepancy '
     'reflects measurement instrument choice.')

Para('Khoảng trống đo lường ở Việt Nam: DISC-5 vs DASS-21 vs RCADS đo các '
     'cấu trúc khác nhau. V-NAMHS dưới ước lượng do ngưỡng lâm sàng; '
     'DASS quá ước lượng do điểm cắt thấp; phiên bản tiếng Việt của RCADS '
     'cung cấp đo lường có tính phát triển phù hợp và đặc thù theo phân '
     'loại — nhưng thiếu DỮ LIỆU CHUẨN cho diễn giải lâm sàng.',
     'Vietnamese measurement gap: DISC-5 vs DASS-21 vs RCADS measure '
     'different constructs. V-NAMHS underestimates due to clinical thresholds; '
     'DASS overestimates due to low cut-offs; the Vietnamese RCADS adaptation '
     'provides developmentally-appropriate subtype-specific measurement — '
     'but lacks NORMATIVE DATA for clinical interpretation.')

Para('Khoảng trống dữ liệu chuẩn mức độ mục: thanh thiếu niên Việt Nam cần '
     'cơ sở đường nền dựa trên dân số (mean/SD theo từng mục) để so sánh '
     'lâm sàng cho từng cá nhân.',
     'Item-level normative data gap: Vietnamese adolescents need population-'
     'based baseline (mean/SD per item) for individual clinical comparison.')


H2('§3. Mục tiêu nghiên cứu và 3 câu hỏi nghiên cứu',
   '§3. Study objectives and 3 Research Questions')

H3('Lý do tiếp cận mô tả (vs suy luận)',
   'Rationale for descriptive (vs inferential) approach')

Bullet('(a) Phát triển công cụ sàng lọc: phân phối mức độ mục xác định '
       'triệu chứng tần suất cao phù hợp cho sàng lọc ngắn',
       '(a) Screening tool development: item-level distributions identify '
       'high-frequency symptoms suitable for brief screening')
Bullet('(b) Dữ liệu chuẩn: nhà lâm sàng Việt Nam cần mean/SD dựa trên dân '
       'số để giải thích điểm cá nhân',
       '(b) Normative data: Vietnamese clinicians need population-based '
       'mean/SD to interpret individual scores')
Bullet('(c) Quỹ đạo phát triển: mẫu hình theo khối lớp thông tin cho thời '
       'điểm can thiệp phù hợp với độ tuổi',
       '(c) Developmental trajectory: grade-level patterns inform age-'
       'appropriate intervention timing')
Bullet('(d) Logic bài đồng hành (Q3-9): phân tích cơ chế tích hợp trong '
       'bài báo đồng hành [Cong và cs., BMC Psychiatry, đang chuẩn bị]',
       '(d) Companion paper logic (Q3-9): integrated mechanism analysis in '
       'companion paper [Cong et al., BMC Psychiatry, in prep]')

H3('Ba câu hỏi nghiên cứu có thể kiểm chứng',
   'Three testable Research Questions')

rqs = [
    (('RQ1', 'RQ1'),
     ('Phân phối mức độ mục của triệu chứng GAD/SAD/SocAD trên mẫu lớn '
      'học sinh THCS Việt Nam là gì, cung cấp dữ liệu đường nền chuẩn?',
      'What is the item-level frequency distribution (M, SD, rank) of '
      'GAD/SAD/SocAD symptoms in a large sample of Vietnamese lower '
      'secondary school students, providing normative baseline data?')),
    (('RQ2', 'RQ2'),
     ('Mức độ triệu chứng RLLA khác biệt theo giới tính và khối lớp (6-9) '
      'như thế nào, và quỹ đạo phát triển của từng dạng là gì?',
      'Do anxiety symptom levels differ by gender and grade level (6-9), '
      'and what is the developmental trajectory of each subtype?')),
    (('RQ3', 'RQ3'),
     ('Mục RCADS nào có điểm trung bình cao nhất, đủ điều kiện làm mục '
      'tiêu sàng lọc ưu tiên cho việc xác định học sinh có rối loạn lo âu '
      'trong bối cảnh trường học Việt Nam?',
      'Which specific RCADS items demonstrate highest mean scores, '
      'qualifying as priority screening targets for adolescent mental '
      'health identification in Vietnamese school context?')),
]
make_table(['Mã', 'Phát biểu câu hỏi nghiên cứu'],
           ['Code', 'Research question statement'],
           rqs, [1.5, 15.5])


d.add_page_break()


# ============================================================
H1('4. PHƯƠNG PHÁP NGHIÊN CỨU', '4. METHODS')


H2('4.1 Thiết kế và đối tượng', '4.1 Design and participants')

Para('Nghiên cứu cắt ngang mô tả. Mẫu N = 1.352 học sinh THCS tại 2 trường '
     'Hà Nội (Nhật Tân + Tây Mỗ). Phân bố giới: 614 nam (45,4%), 738 nữ '
     '(54,6%). Phân bố khối: 6 (368), 7 (316), 8 (340), 9 (328). Độ tuổi '
     '11-14. Chiến lược chọn mẫu cụm 2 giai đoạn: trường được chọn mục '
     'đích cho đại diện nội thành/ngoại thành; tất cả học sinh trong các '
     'lớp được chọn được mời tham gia.',
     'Cross-sectional descriptive study. Sample N = 1,352 lower secondary '
     'school students at 2 schools in Hanoi (Nhat Tan + Tay Mo). Gender '
     'distribution: 614 male (45.4%), 738 female (54.6%). Grade distribution: '
     '6 (368), 7 (316), 8 (340), 9 (328). Age 11-14. 2-stage cluster sampling: '
     'schools selected purposively for urban/suburban representation; all '
     'students in selected classes invited to participate.')


H2('4.2 Công cụ: phiên bản tiếng Việt của RCADS',
   '4.2 Instrument: Vietnamese RCADS adaptation')

Para('RCADS (Revised Children\'s Anxiety and Depression Scale; Chorpita, '
     '2000) — tiêu chuẩn vàng đo các phân loại RLLA theo DSM-5 cho thanh '
     'thiếu niên. Quy trình thích nghi: dịch xuôi-dịch ngược; đánh giá '
     'chuyên gia (3 nhà tâm lý trẻ em + 2 nhà giáo dục); khảo sát thử '
     'nghiệm (n = 50). Từ 21 mục gốc, giữ lại 15 mục cho 3 phân loại:',
     'RCADS (Revised Children\'s Anxiety and Depression Scale; Chorpita, '
     '2000) — gold standard for DSM-5 subtype-specific adolescent anxiety. '
     'Adaptation: forward-back translation; expert review (3 child '
     'psychologists + 2 educators); pilot testing (n = 50). From 21 original '
     'items, 15 items retained across 3 subtypes:')

Bullet('GAD (7 mục): RCADS1, 4, 8, 12, 13, 30, 35',
       'GAD (7 items): RCADS1, 4, 8, 12, 13, 30, 35')
Bullet('SAD (4 mục): RCADS5, 17, 45, 46', 'SAD (4 items): RCADS5, 17, 45, 46')
Bullet('SocAD (4 mục): RCADS20, 32, 38, 43',
       'SocAD (4 items): RCADS20, 32, 38, 43')

Para('Thang Likert 4 mức: 0 = Không bao giờ, 1 = Đôi khi, 2 = Thường xuyên, '
     '3 = Luôn luôn. Điểm thô được chuyển đổi sang thang 0-100 cho so '
     'sánh chuẩn hóa.',
     '4-point Likert: 0 = Never, 1 = Sometimes, 2 = Often, 3 = Always. Raw '
     'scores converted to 0-100 scale for standardized comparison.')


H2('4.3 Độ tin cậy', '4.3 Reliability')

Para('Cronbach α và McDonald ω được tính cho từng dạng RLLA. Tiêu chí: '
     'α ≥ 0,70 cho độ tin cậy chấp nhận (Nunnally, 1978).',
     'Cronbach α and McDonald ω computed per anxiety subtype. Criterion: '
     'α ≥ 0.70 for adequate reliability (Nunnally, 1978).')


H2('4.4 Chiến lược phân tích', '4.4 Analytic strategy')

H3('Thống kê mô tả (RQ1)', 'Descriptive statistics (RQ1)')
Bullet('Trung bình, SD, range cho từng mục',
       'Mean, SD, range per item')
Bullet('Xếp hạng mục trong từng phân loại (1 = M cao nhất → thấp nhất)',
       'Within-subtype item ranking (1 = highest M to lowest)')
Bullet('Điểm trung bình tổng theo từng phân loại',
       'Subtype aggregate scores')

H3('Thống kê suy luận với hiệu chỉnh (RQ2)',
   'Inferential statistics with corrections (RQ2)')
Bullet('So sánh giới: kiểm định t độc lập hoặc ANOVA một chiều với HIỆU '
       'CHỈNH BONFERRONI (α = 0,05/3 = 0,0167 cho 3 phân loại)',
       'Gender comparison: independent t-test or one-way ANOVA with '
       'BONFERRONI CORRECTION (α = 0.05/3 = 0.0167 for 3 subtypes)')
Bullet('So sánh khối lớp: ANOVA một chiều + hậu kiểm Tukey HSD',
       'Grade comparison: one-way ANOVA + post-hoc Tukey HSD')
Bullet('Cỡ tác động: Cohen d cho so sánh giới (nhỏ = 0,2; trung bình = 0,5; '
       'lớn = 0,8); η² cho so sánh khối lớp (nhỏ = 0,01; trung bình = 0,06; '
       'lớn = 0,14)',
       'Effect sizes: Cohen d for gender (small = 0.2, medium = 0.5, '
       'large = 0.8); partial η² for grade (small = 0.01, medium = 0.06, '
       'large = 0.14)')
Bullet('Khoảng tin cậy 95% cho tất cả các giá trị trung bình',
       '95% Confidence Intervals for all means')

H3('Xác định mục tiêu sàng lọc (RQ3)', 'Screening target identification (RQ3)')
Bullet('Mục có M ≥ 50,0 (trên giữa thang 0-100) đánh dấu là mục tiêu sàng '
       'lọc ưu tiên',
       'Items with M ≥ 50.0 (above middle of 0-100 scale) flagged as priority '
       'screening targets')
Bullet('Mục có SD ≤ 30 chỉ ra triệu chứng tần suất cao nhất quán (công cụ '
       'sàng lọc đáng tin cậy hơn)',
       'Items with SD ≤ 30 indicate consistent high-frequency symptoms '
       '(more reliable screeners)')


H2('4.5 Đạo đức nghiên cứu — đầy đủ (Q3-6 fix)',
   '4.5 Ethics — full statement (Q3-6 fix)')

NCS_marker('NCS bổ sung số quyết định + ngày + tên Hội đồng Đạo đức HNUE.')

Para('Tuyên bố đầy đủ tạm thời (chờ NCS xác nhận): "Nghiên cứu này đã '
     'được Hội đồng Đạo đức Nghiên cứu của Trường Đại học Sư phạm Hà Nội '
     '(HNUE) phê duyệt vào ngày [...] (số phê duyệt: [...]). Đã có sự '
     'đồng ý bằng văn bản của cha mẹ/người giám hộ hợp pháp và sự đồng '
     'thuận bằng văn bản của tất cả học sinh tham gia. Dữ liệu được mã hóa '
     'ẩn danh bằng các định danh duy nhất, không có thông tin cá nhân nào '
     'liên kết với phản hồi của từng cá nhân. Dữ liệu được lưu trữ trên '
     'máy chủ tổ chức được bảo vệ mật khẩu, chỉ có nhóm nghiên cứu mới được '
     'truy cập. Quy trình nghiên cứu tuân thủ Tuyên bố Helsinki (2013) và '
     'Luật Trẻ em Việt Nam (2016)."',
     'Tentative full statement (pending NCS confirmation): "This study was '
     'approved by the Institutional Review Board of Hanoi National University '
     'of Education (HNUE) on [date] (approval number: [...]). Written '
     'informed consent was obtained from parents/legal guardians, and '
     'written assent was obtained from all student participants. Data were '
     'anonymized using unique identifiers, with no personal identifiers '
     'linked to individual responses. Data are stored on password-protected '
     'institutional servers, accessible only to the research team. Study '
     'procedures complied with the Declaration of Helsinki (2013) and '
     'Vietnamese Law on Children (2016)."')


d.add_page_break()


# ============================================================
H1('5. KẾT QUẢ DỰ KIẾN', '5. EXPECTED RESULTS')


H2('5.1 RQ1 — Phân phối mức độ mục', '5.1 RQ1 — Item-level distribution')

H3('5.1.1 Bảng 2: Rối loạn lo âu lan toả (GAD)',
   '5.1.1 Table 2: Generalized Anxiety Disorder (GAD)')

gad_data = [
    (('1', '1'),
     ('RCADS4', 'RCADS4'),
     ('Tôi lo lắng khi nghĩ rằng mình đã không làm tốt điều gì đó',
      'I worry when I think I have done poorly at something'),
     ('64,28', '64.28'),
     ('29,69', '29.69')),
    (('2', '2'),
     ('RCADS13', 'RCADS13'),
     ('Tôi lo lắng rằng điều gì đó tồi tệ sẽ xảy ra',
      'I worry that something bad will happen'),
     ('59,62', '59.62'),
     ('35,86', '35.86')),
    (('3', '3'),
     ('RCADS8', 'RCADS8'),
     ('Tôi cảm thấy lo lắng khi nghĩ rằng ai đó...',
      'I feel worried when I think someone...'),
     ('59,02', '59.02'),
     ('33,48', '33.48')),
    (('4', '4'),
     ('RCADS30', 'RCADS30'),
     ('Tôi lo lắng về việc mắc lỗi',
      'I worry about making mistakes'),
     ('57,69', '57.69'),
     ('31,26', '31.26')),
    (('5', '5'),
     ('RCADS12', 'RCADS12'),
     ('Tôi lo lắng rằng mình sẽ làm bài tập ở t...',
      'I worry that I will do badly at school work'),
     ('55,13', '55.13'),
     ('33,84', '33.84')),
    (('6', '6'),
     ('RCADS1', 'RCADS1'),
     ('Tôi lo lắng về mọi thứ', 'I worry about everything'),
     ('49,14', '49.14'),
     ('30,72', '30.72')),
    (('7', '7'),
     ('RCADS35', 'RCADS35'),
     ('Tôi lo lắng về những gì sẽ xảy ra',
      'I worry about what is going to happen'),
     ('45,86', '45.86'),
     ('33,83', '33.83')),
]
make_table(['Xếp hạng', 'Mã mục', 'Nội dung mục', 'M (0-100)', 'SD'],
           ['Rank', 'Item code', 'Item content', 'M (0-100)', 'SD'],
           gad_data, [1.5, 2.0, 7.5, 2.5, 2.5])

Para('Mức trung bình GAD: M = 55,82 — mức trung bình cao.',
     'GAD subtype mean: M = 55.82 — moderate-high level.', indent=False)


H3('5.1.2 Bảng 3: Rối loạn lo âu chia ly (SAD)',
   '5.1.2 Table 3: Separation Anxiety Disorder (SAD)')

sad_data = [
    (('1', '1'),
     ('RCADS46', 'RCADS46'),
     ('Tôi sẽ cảm thấy sợ nếu phải ở xa nhà qua đêm',
      'I would feel afraid of being on my own at night'),
     ('27,88', '27.88')),
    (('2', '2'),
     ('RCADS17', 'RCADS17'),
     ('Tôi cảm thấy sợ nếu phải ngủ một mình',
      'I feel scared if I have to sleep on my own'),
     ('25,49', '25.49')),
    (('3', '3'),
     ('RCADS5', 'RCADS5'),
     ('Tôi cảm thấy sợ khi ở một mình ở nhà',
      'I feel scared when I have to stay on my own'),
     ('25,35', '25.35')),
    (('4', '4'),
     ('RCADS45', 'RCADS45'),
     ('Tôi lo lắng khi đi ngủ vào ban đêm',
      'I worry about going to bed at night'),
     ('21,52', '21.52')),
]
make_table(['Xếp hạng', 'Mã mục', 'Nội dung mục', 'M (0-100)'],
           ['Rank', 'Item code', 'Item content', 'M (0-100)'],
           sad_data, [1.5, 2.0, 9.0, 4.5])

Para('Mức trung bình SAD: M = 25,06 — THẤP NHẤT trong 3 dạng, phù hợp với '
     'phân loại DSM-5: SAD = khởi phát từ thời thơ ấu, giảm dần qua tuổi '
     'vị thành niên.',
     'SAD subtype mean: M = 25.06 — LOWEST of 3 subtypes, consistent with '
     'DSM-5 classification: SAD = childhood-onset, declining through '
     'adolescence.', indent=False)


H3('5.1.3 Bảng 4: Rối loạn lo âu xã hội (SocAD)',
   '5.1.3 Table 4: Social Anxiety Disorder (SocAD)')

socad_data = [
    (('1', '1'),
     ('RCADS32', 'RCADS32'),
     ('Tôi lo lắng về việc người khác nghĩ gì về mình',
      'I worry about what other people think of me'),
     ('56,98', '56.98')),
    (('2', '2'),
     ('RCADS43', 'RCADS43'),
     ('Tôi sợ rằng mình sẽ làm trò cười trước mọi người',
      'I am afraid that I will make a fool of myself'),
     ('49,26', '49.26')),
    (('3', '3'),
     ('RCADS38', 'RCADS38'),
     ('Tôi cảm thấy sợ nếu phải nói trước lớp',
      'I feel scared when I have to talk in front of my class'),
     ('45,32', '45.32')),
    (('4', '4'),
     ('RCADS20', 'RCADS20'),
     ('Tôi lo rằng mình sẽ trông ngốc nghếch',
      'I worry that I will look stupid'),
     ('42,09', '42.09')),
]
make_table(['Xếp hạng', 'Mã mục', 'Nội dung mục', 'M (0-100)'],
           ['Rank', 'Item code', 'Item content', 'M (0-100)'],
           socad_data, [1.5, 2.0, 9.0, 4.5])

Para('Mức trung bình SocAD: M = 48,41 — mức trung bình.',
     'SocAD subtype mean: M = 48.41 — moderate level.', indent=False)


H2('5.2 RQ2 — Khác biệt nhân khẩu', '5.2 RQ2 — Demographic differences')

H3('5.2.1 So sánh giới với hiệu chỉnh Bonferroni',
   '5.2.1 Gender comparison with Bonferroni correction')

gender_q3 = [
    (('GAD', 'GAD'),
     ('51,43 (22,01)', '51.43 (22.01)'),
     ('59,47 (22,07)', '59.47 (22.07)'),
     ('44,484', '44.484'),
     ('< 0,001', '< 0.001'),
     ('0,37', '0.37')),
    (('SAD ⭐', 'SAD ⭐'),
     ('25,42 (25,46)', '25.42 (25.46)'),
     ('24,76 (23,29)', '24.76 (23.29)'),
     ('0,246', '0.246'),
     ('0,620', '0.620'),
     ('0,03', '0.03')),
    (('SocAD', 'SocAD'),
     ('43,20 (25,09)', '43.20 (25.09)'),
     ('52,74 (26,31)', '52.74 (26.31)'),
     ('45,984', '45.984'),
     ('< 0,001', '< 0.001'),
     ('0,37', '0.37')),
]
make_table(['Dạng', 'Nam M (SD)', 'Nữ M (SD)', 'F', 'p', 'Cohen d'],
           ['Type', 'Male M (SD)', 'Female M (SD)', 'F', 'p', 'Cohen d'],
           gender_q3, [2.0, 3.0, 3.0, 2.5, 2.5, 2.0])

Para('Lưu ý ngưỡng có ý nghĩa hiệu chỉnh Bonferroni: α = 0,0167.',
     'Note Bonferroni-corrected significance threshold: α = 0.0167.',
     indent=False, italic=True)

Para('⭐ SAD không có ý nghĩa thống kê khác biệt giới — phát hiện này được '
     'phân tích chi tiết qua kiểm định bất biến đa nhóm trong bài đồng hành.',
     '⭐ SAD shows no statistically significant gender difference — this '
     'finding is analyzed in detail through multi-group invariance testing '
     'in the companion paper.', indent=False, italic=True)


H3('5.2.2 So sánh khối lớp + hậu kiểm Tukey',
   '5.2.2 Grade comparison + post-hoc Tukey')

grade_q3 = [
    (('GAD', 'GAD'),
     ('54,32 → 53,65 → 55,63 → 59,79',
      '54.32 → 53.65 → 55.63 → 59.79'),
     ('5,020', '5.020'),
     ('0,002', '0.002'),
     ('0,011 nhỏ', '0.011 small'),
     ('Khối 9 > 6, 7', 'Grade 9 > 6, 7')),
    (('SAD ⭐', 'SAD ⭐'),
     ('32,13 → 27,14 → 20,88 → 19,46',
      '32.13 → 27.14 → 20.88 → 19.46'),
     ('21,239', '21.239'),
     ('< 0,001', '< 0.001'),
     ('0,045 vừa', '0.045 medium'),
     ('GIẢM ĐƠN ĐIỆU theo khối lớp',
      'MONOTONIC DECLINE across grades')),
    (('SocAD', 'SocAD'),
     ('Mẫu hình thay đổi nhẹ',
      'Slight pattern change'),
     ('4,879', '4.879'),
     ('0,002', '0.002'),
     ('0,011', '0.011'),
     ('Khác biệt giữa các khối',
      'Inter-grade differences')),
]
make_table(['Dạng', 'Quỹ đạo khối 6 → 9', 'F', 'p', 'η²', 'Tukey'],
           ['Type', 'Grade 6 → 9 trajectory', 'F', 'p', 'η²', 'Tukey'],
           grade_q3, [2.0, 4.5, 2.0, 2.0, 2.5, 4.0])

Para('⭐ Quỹ đạo SAD giảm đơn điệu (32,13 → 19,46) — phù hợp DSM-5 SAD = '
     'khởi phát thời thơ ấu giảm theo tuổi.',
     '⭐ SAD trajectory shows monotonic decline (32.13 → 19.46) — consistent '
     'with DSM-5 SAD = childhood-onset declining with age.',
     indent=False, italic=True)


H2('5.3 RQ3 — Mục tiêu sàng lọc ưu tiên (Bảng 6)',
   '5.3 RQ3 — Priority screening targets (Table 6)')

screen = [
    (('Sàng lọc GAD (3 mục)', 'GAD screener (3 items)'),
     ('RCADS4 + RCADS13 + RCADS8',
      'RCADS4 + RCADS13 + RCADS8'),
     ('60,97', '60.97'),
     ('Sàng lọc lo âu hiệu suất + tổng quát',
      'Performance and general anxiety screening')),
    (('Sàng lọc SocAD (2 mục)', 'SocAD screener (2 items)'),
     ('RCADS32 + RCADS43',
      'RCADS32 + RCADS43'),
     ('53,12', '53.12'),
     ('Sàng lọc nỗi sợ phán xét xã hội',
      'Social judgment fear screening')),
    (('Đánh giá SAD (2 mục)', 'SAD assessment (2 items)'),
     ('RCADS46 + RCADS17',
      'RCADS46 + RCADS17'),
     ('< 30,0', '< 30.0'),
     ('Không đạt ngưỡng sàng lọc — chuyển sang đánh giá lâm sàng',
      'Below screening threshold — clinical assessment recommended')),
]
make_table(['Công cụ sàng lọc', 'Mục đề xuất', 'M kết hợp', 'Mục đích'],
           ['Screener', 'Recommended items', 'Combined M', 'Purpose'],
           screen, [4.0, 4.5, 2.5, 6.0])


H2('5.4 Hình 1 — Biểu đồ quỹ đạo theo khối lớp',
   '5.4 Figure 1 — Grade trajectory visualization')

Bullet('Biểu đồ đường thể hiện thay đổi điểm trung bình theo 4 khối lớp',
       'Line chart showing mean score changes across 4 grades')
Bullet('GAD: ổn định rồi tăng (53,65 → 59,79)',
       'GAD: stable then increase (53.65 → 59.79)')
Bullet('SAD: GIẢM đơn điệu (32,13 → 19,46) — phát hiện trực quan then chốt',
       'SAD: monotonic DECLINE (32.13 → 19.46) — key visualization')
Bullet('SocAD: tăng nhẹ',
       'SocAD: subtle increase')


d.add_page_break()


# ============================================================
H1('6. PHẦN BÀN LUẬN', '6. DISCUSSION')


H2('6.1 Tóm tắt phát hiện chính', '6.1 Summary of key findings')

Bullet('RQ1: Dữ liệu mức độ mục cung cấp đường nền chuẩn đầu tiên cho Việt '
       'Nam; GAD chủ yếu là triệu chứng lo âu thành tích học tập; SocAD '
       'chủ yếu là triệu chứng sợ phán xét xã hội; SAD chủ yếu là các mục '
       'liên quan đến ban đêm/xa nhà',
       'RQ1: Item-level data provides first Vietnamese normative baseline; '
       'GAD predominantly academic-worry symptoms; SocAD predominantly '
       'social-judgment symptoms; SAD predominantly nighttime/away-from-home '
       'items')
Bullet('RQ2: GAD và SocAD cao hơn rõ rệt ở nữ; SAD KHÔNG có khác biệt giới; '
       'SAD giảm đơn điệu theo khối lớp (phát triển); GAD tăng theo khối lớp',
       'RQ2: GAD and SocAD significantly higher in females; SAD shows NO '
       'gender difference; SAD declines monotonically with grade '
       '(developmental); GAD increases with grade')
Bullet('RQ3: 5 mục sàng lọc ưu tiên được xác định trên 2 phân loại',
       'RQ3: 5 priority screening items identified across 2 subtypes')


H2('6.2 Hàm ý phát triển', '6.2 Developmental implications')

Para('Quỹ đạo SAD giảm đơn điệu (η² = 0,045 cỡ tác động vừa) ủng hộ phân '
     'loại DSM-5 và ICD-11 về SAD như là rối loạn chủ yếu khởi phát thời '
     'thơ ấu, với thanh thiếu niên đang chuyển ra khỏi lo âu dựa trên gắn '
     'bó. GAD tăng nhẹ qua các khối — phù hợp với y văn phát triển về '
     'GAD khởi phát muộn-thơ ấu/vị thành niên (Beesdo-Baum và cs., 2022).',
     'SAD monotonic decline (η² = 0.045 medium effect) supports DSM-5 ICD-11 '
     'classification of SAD as primarily childhood-onset, with adolescents '
     'transitioning out of attachment-driven anxiety. GAD slight increase '
     'across grades — consistent with developmental literature on '
     'late-childhood/adolescent GAD onset (Beesdo-Baum et al. 2022).')


H2('6.3 Mẫu hình giới (Q3-8 fix — ngắn gọn nhưng có cơ sở)',
   '6.3 Gender pattern observations (Q3-8 fix — brief but justified)')

Para('Nữ > nam ở GAD và SocAD phù hợp với y văn quốc tế (McLean và cs., '
     '2011). TÍNH BẤT BIẾN GIỚI của SAD — phù hợp với bằng chứng quốc tế '
     'cho lứa tuổi nhỏ nhưng MỚI cho thanh thiếu niên Việt Nam.',
     'Female > male in GAD and SocAD consistent with international literature '
     '(McLean et al., 2011). SAD GENDER INVARIANCE — consistent with '
     'international evidence for younger ages but NOVEL for Vietnamese '
     'adolescents.')

NCS_marker('Quyết định Q3-9 cần thiết để hoàn thiện chiến lược citation Q1 '
           '↔ Q3.')

Para('Tuyên bố tạm thời: "Phân tích cơ chế chi tiết của các mẫu hình giới '
     'khác biệt giữa các phân loại, bao gồm kiểm định bất biến SEM đa nhóm, '
     'được trình bày trong bài phân tích cơ chế đồng hành [Cong và cs., '
     'đang xét duyệt tại BMC Psychiatry]."',
     'Tentative statement: "Detailed mechanism analysis of differential '
     'gender patterns across subtypes, including multi-group structural '
     'equation modeling invariance testing, is presented in our companion '
     'mechanistic analysis [Cong et al., under review at BMC Psychiatry]."')


H2('6.4 Hàm ý thực hành', '6.4 Practical implications')

Bullet('Công cụ sàng lọc 5 mục khả thi từ các mục đứng đầu mỗi phân loại — '
       'có thể tích hợp vào sàng lọc trường học toàn cầu',
       '5-item screening tool feasible from top items per subtype — could be '
       'integrated into school-based universal screening')
Bullet('Điểm cắt: các mục có M > 60 trong mẫu này có thể được dùng làm '
       'tham chiếu chuẩn sơ bộ',
       'Cut-points: items with M > 60 in this sample may serve as preliminary '
       'normative reference')
Bullet('Can thiệp mục tiêu theo khối lớp: ưu tiên sàng lọc SAD khối 6-7; '
       'ưu tiên sàng lọc GAD/SocAD khối 8-9',
       'Grade-targeted intervention: prioritize SAD screening grades 6-7; '
       'prioritize GAD/SocAD screening grades 8-9')


H2('6.5 Giới hạn và hướng nghiên cứu tương lai',
   '6.5 Limitations and future research')

Bullet('Thiết kế cắt ngang — không thể suy luận nhân quả',
       'Cross-sectional design — cannot infer causation')
Bullet('2 trường Hà Nội → chỉ nội thành/ngoại thành; cần lặp lại ở nông thôn',
       '2 schools Hanoi → urban/suburban only; need rural replication')
Bullet('Sai lệch tự báo cáo cho triệu chứng nội tâm hóa',
       'Self-report bias for internalizing symptoms')
Bullet('Phiên bản tiếng Việt RCADS — cần các nghiên cứu xác thực thêm '
       '(tính hợp lệ đồng thời với chẩn đoán lâm sàng)',
       'Vietnamese RCADS adaptation — needs further validation studies '
       '(concurrent validity with clinician diagnosis)')
Bullet('Hướng tương lai: nghiên cứu theo dõi quỹ đạo theo khối lớp (theo '
       'dõi đoàn hệ); phân tích ROC cho điểm cắt sàng lọc; so sánh chéo '
       'văn hóa',
       'Future: longitudinal grade trajectory (cohort follow-up); ROC '
       'analysis for screening cut-points; cross-cultural comparison')


d.add_page_break()


# ============================================================
H1('7. TÀI LIỆU THAM KHẢO DỰ KIẾN', '7. EXPECTED REFERENCES')

Para('Dự kiến 30-35 tài liệu tham khảo. Tất cả đã được kiểm chứng đối với '
     'PDF gốc trong kho dữ liệu.',
     'Approximately 30-35 references. All verified against original PDFs in '
     'the research database.', indent=False)


H2('7.1 Nghiên cứu thực nghiệm (8 PDFs đã kiểm chứng)',
   '7.1 Empirical studies (8 verified PDFs)')

Bullet('Xu et al. (2021), Wen et al. (2020), Anderson et al. (2025), V-NAMHS '
       '(2022), Hoang Trung Hoc et al. (2025), Bhardwaj et al. (2020), '
       'Nakie et al. (2022), Saikia et al. (2023)',
       'Xu et al. (2021), Wen et al. (2020), Anderson et al. (2025), V-NAMHS '
       '(2022), Hoang Trung Hoc et al. (2025), Bhardwaj et al. (2020), '
       'Nakie et al. (2022), Saikia et al. (2023)')


H2('7.2 Phương pháp luận và lý thuyết', '7.2 Methodological and theoretical')

Bullet('Chorpita (2000) — phát triển thang RCADS',
       'Chorpita (2000) — RCADS scale development')
Bullet('Rosenberg (1965) — lý thuyết chuẩn tâm trắc',
       'Rosenberg (1965) — psychometric norm theory')
Bullet('Nunnally (1978) — độ tin cậy α',
       'Nunnally (1978) — α reliability')
Bullet('Beesdo-Baum et al. (2022) — phát triển GAD ở vị thành niên',
       'Beesdo-Baum et al. (2022) — adolescent GAD development')
Bullet('McLean et al. (2011) — phân tích meta khác biệt giới ở rối loạn lo âu',
       'McLean et al. (2011) — gender differences in anxiety meta-analysis')
Bullet('GBD ASEAN (2025) — gánh nặng bệnh tật khu vực',
       'GBD ASEAN (2025) — regional disease burden')


H2('7.3 Bài đồng hành (Q3-9 BLOCKING)',
   '7.3 Companion paper (Q3-9 BLOCKING)')

NCS_marker('Thầy + NCS quyết chiến lược citation: (A) Q1 chấp nhận trước → '
           'Q3 cite Q1, HAY (B) Q1 + Q3 nộp cùng lúc với chú thích '
           '"under review"')


# ============================================================
d.add_page_break()
H1('8. KẾ HOẠCH BẢNG VÀ HÌNH', '8. TABLES AND FIGURES PLAN')

t_f = [
    (('Bảng 1', 'Table 1'),
     ('Đặc điểm mẫu', 'Sample demographics'),
     ('Phân bố theo giới × khối × trường',
      'Distribution by gender × grade × school')),
    (('Bảng 2', 'Table 2'),
     ('Mức độ mục GAD', 'GAD item-level'),
     ('7 mục với M, SD, xếp hạng', '7 items with M, SD, rank')),
    (('Bảng 3', 'Table 3'),
     ('Mức độ mục SAD', 'SAD item-level'),
     ('4 mục với M, SD, xếp hạng', '4 items with M, SD, rank')),
    (('Bảng 4', 'Table 4'),
     ('Mức độ mục SocAD', 'SocAD item-level'),
     ('4 mục với M, SD, xếp hạng', '4 items with M, SD, rank')),
    (('Bảng 5a', 'Table 5a'),
     ('So sánh giới với Cohen d',
      'Gender comparison with Cohen d'),
     ('F, p với hiệu chỉnh Bonferroni',
      'F, p with Bonferroni correction')),
    (('Bảng 5b', 'Table 5b'),
     ('So sánh khối lớp với η²', 'Grade comparison with η²'),
     ('Hậu kiểm Tukey HSD', 'Post-hoc Tukey HSD')),
    (('Bảng 6', 'Table 6'),
     ('TÓM TẮT mục tiêu sàng lọc', 'SUMMARY screening targets'),
     ('Q3-7 fix — sàng lọc GAD + SocAD',
      'Q3-7 fix — GAD + SocAD screeners')),
    (('Hình 1', 'Figure 1'),
     ('Biểu đồ quỹ đạo theo khối lớp',
      'Grade trajectory plot'),
     ('Q3-7 fix — TRỰC QUAN THEN CHỐT của bài',
      'Q3-7 fix — KEY VISUALIZATION of paper')),
]
make_table(['Mã', 'Tiêu đề', 'Mô tả nội dung'],
           ['Code', 'Title', 'Content description'],
           t_f, [1.5, 4.5, 11.0])


# ============================================================
H1('9. CÂU HỎI CHỜ QUYẾT (BLOCKING)', '9. BLOCKING QUESTIONS')

Para('Hai câu hỏi sau cần thầy hướng dẫn và NCS quyết định trước khi viết '
     'bản thảo:',
     'The following two questions require decisions from supervisor and PhD '
     'candidate before drafting:', indent=False)

blocking_q3 = [
    (('Q3-6', 'Q3-6'),
     ('Văn bản phê duyệt đạo đức HNUE', 'HNUE ethics approval letter'),
     ('NCS', 'NCS'),
     ('Đã có văn bản chính thức chưa? Nếu chưa, xin được retroactive '
      '(2-4 tuần)?',
      'Is there a formal letter? If not, retroactive approval possible '
      '(2-4 weeks)?')),
    (('Q3-9', 'Q3-9'),
     ('Chiến lược tham chiếu chéo Q1 ↔ Q3',
      'Q1 ↔ Q3 cross-reference strategy'),
     ('Thầy + NCS', 'Supervisor + NCS'),
     ('(A) Nộp Q1 trước (BMC Psychiatry), chờ chấp nhận, nộp Q3 với citation '
      'Q1, HAY (B) nộp Q1 và Q3 cùng lúc với chú thích "companion paper '
      'under review"',
      '(A) Submit Q1 first (BMC Psychiatry), wait for acceptance, submit Q3 '
      'with Q1 citation, OR (B) Submit Q1 and Q3 simultaneously with '
      '"companion paper under review" note')),
]
make_table(['Mã', 'Vấn đề', 'Người quyết', 'Chi tiết quyết định'],
           ['Code', 'Issue', 'Decision maker', 'Decision details'],
           blocking_q3, [1.5, 4.0, 2.5, 9.0])


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
