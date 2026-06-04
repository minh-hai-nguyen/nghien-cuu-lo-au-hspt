# -*- coding: utf-8 -*-
"""Generate v3: KTC vs CrI — chi tiết + dễ hiểu (storytelling + analogies + step-by-step)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Giai_thich_KTC_vs_CrI_cho_thay_v3.docx')

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def PRun(parts):
    p = doc.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 12))
        if opts.get('bold'): r.bold = True
        if opts.get('italic'): r.italic = True
        if opts.get('red'): r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif opts.get('color'): r.font.color.rgb = opts['color']
    return p

def H1(t): return P(t, bold=True, size=16, color=RGBColor(0x1F, 0x3A, 0x68), align=WD_ALIGN_PARAGRAPH.CENTER)
def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10.5)
    return t

def note_box(text, color_hex='FFF3E0'):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    c.text = ''
    r = c.paragraphs[0].add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.italic = True
    set_cell_bg(c, color_hex)


# ============================================================
# TIÊU ĐỀ
# ============================================================
H1('HIỂU VỀ KTC VÀ CrI TRONG 15 PHÚT')
H1('Dành cho người làm lâm sàng/nghiên cứu tâm lý học')
P('Giải thích từ bối cảnh thực tế — không đi vào công thức phức tạp',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=RGBColor(0x66, 0x66, 0x66))
P('Nhóm dự án Lo âu | Tháng 04/2026',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P()

# ============================================================
# PHẦN 0 — TÓM TẮT 30 GIÂY
# ============================================================
note_box(
    'TÓM TẮT TRONG 30 GIÂY (cho người bận):\n\n'
    'Khi đọc báo cáo nghiên cứu có ghi "SMD = –0,19 (95 % CrI: –0,22; –0,17)", thầy có thể '
    'đọc nó chính xác như một khoảng tin cậy 95 % thông thường (CI). Sự khác biệt giữa CI và CrI chủ '
    'yếu về TRIẾT HỌC thống kê (tần suất vs Bayesian), KHÔNG ảnh hưởng kết luận lâm sàng. '
    'Trong 90 % trường hợp thực tế tại tâm lý học, CI và CrI cho KẾT QUẢ SỐ gần như trùng nhau. '
    'Điều thực sự quan trọng khi đọc: (1) Khoảng có chứa 0 không? → có ý nghĩa thống kê không; '
    '(2) Độ lớn effect size? → có ý nghĩa lâm sàng không.',
    color_hex='FFFDE7')
P()

# ============================================================
# PHẦN 1 — CÂU CHUYỆN MỞ ĐẦU
# ============================================================
H2('1. Câu chuyện mở đầu: Tại sao lại có hai khoảng tin cậy khác nhau?')

P('Hãy tưởng tượng thầy có một học viên hỏi:', italic=True)
P()
P('"Thưa cô, em làm nghiên cứu về stress ở học sinh. Trong mẫu 100 em, điểm stress trung bình '
  'là 15,2 ± 3,4 (SD). Cô cho em hỏi: điểm stress THỰC SỰ của QUẦN THỂ học sinh Việt Nam là bao nhiêu?"',
  italic=True, color=RGBColor(0x2E, 0x54, 0x8B))
P()
P('Có HAI cách trả lời, tuỳ vào trường phái thống kê:', bold=True)
P()

P('Cách 1 — Trường phái TẦN SUẤT (Frequentist, chuẩn cổ điển):', bold=True)
P('"Từ mẫu của em, tôi tính được khoảng tin cậy 95 %: 14,5 đến 15,9. Ý nghĩa: NẾU em lặp lại '
  'nghiên cứu này 100 lần (mỗi lần mẫu 100 em khác nhau), thì khoảng 95 trong số 100 khoảng '
  'tin cậy tính được sẽ CHỨA giá trị thực. Nhưng cụ thể khoảng [14,5; 15,9] này CÓ CHỨA giá '
  'trị thực hay không — tôi KHÔNG biết, vì giá trị thực chỉ có thể là MỘT SỐ cố định, còn '
  'khoảng tôi tính là ngẫu nhiên."')
PRun([
    ('→ Cái này gọi là ', {}),
    ('CI', {'bold': True}),
    (' = ', {}),
    ('C', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('onfidence ', {}),
    ('I', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('nterval.', {}),
])
P()

P('Cách 2 — Trường phái BAYESIAN:', bold=True)
P('"Từ mẫu của em + niềm tin ban đầu hợp lý (prior), tôi tính được khoảng credible 95 %: '
  '14,5 đến 15,9. Ý nghĩa: CÓ 95 % XÁC SUẤT rằng giá trị thực nằm trong khoảng [14,5; 15,9]. '
  'Đơn giản vậy thôi."')
PRun([
    ('→ Cái này gọi là ', {}),
    ('CrI', {'bold': True}),
    (' = ', {}),
    ('Cr', {'bold': True, 'color': RGBColor(0x0B, 0x5E, 0x20)}),
    ('edible ', {}),
    ('I', {'bold': True, 'color': RGBColor(0x0B, 0x5E, 0x20)}),
    ('nterval.', {}),
])
P()

note_box(
    'NHẬN XÉT: Hai khoảng có thể TRÙNG HOÀN TOÀN về số (ví dụ cả hai đều [14,5; 15,9]) nhưng '
    'DIỄN GIẢI khác nhau. CI là về "procedure" (phương pháp — nếu lặp lại nhiều lần), CrI là '
    'về "probability" (xác suất — cho khoảng cụ thể này). Đây là vấn đề TRIẾT HỌC thống kê, '
    'không phải vấn đề tính toán.',
    color_hex='E3F2FD')
P()

# ============================================================
# PHẦN 2 — PHÉP ẨN DỤ
# ============================================================
H2('2. Phép ẩn dụ giúp nhớ: Câu cá vs Bẫy cá')

P('Hãy tưởng tượng thầy muốn biết con cá LỚN NHẤT trong ao đang ở ĐÂU:', italic=True)
P()

H3('CI — như "THẢ LƯỚI nhiều lần" (Frequentist)')
P('"Tôi thả 100 cái lưới (mỗi lần thả là một nghiên cứu mới). 95 trong 100 lưới sẽ BẮT được '
  'con cá lớn nhất. Nhưng CÁI LƯỚI THẦY ĐANG CẦM — tôi không biết nó có bắt được cá hay không. '
  'Tôi chỉ biết TỶ LỆ THÀNH CÔNG của phương pháp thả lưới."', italic=True)
PRun([
    ('→ ', {}),
    ('Tuyên bố về phương pháp', {'bold': True}),
    (', không phải về lưới cụ thể.', {}),
])
P()

H3('CrI — như "NIỀM TIN VỀ VỊ TRÍ con cá" (Bayesian)')
P('"Dựa trên quan sát + kinh nghiệm trước, TÔI TIN RẰNG có 95 % khả năng con cá lớn nhất đang '
  'nằm trong góc AO mà tôi đã khoanh. Nói ngắn gọn: có 95 % xác suất cá ở khoảng này."', italic=True)
PRun([
    ('→ ', {}),
    ('Tuyên bố trực tiếp về khoảng cụ thể', {'bold': True}),
    ('.', {}),
])
P()

note_box(
    'ĐIỂM QUAN TRỌNG: Đa số người đọc bài báo (cả thầy, học viên, bác sĩ) ĐỀU hiểu CI '
    'như thể là CrI — tức "95 % xác suất giá trị thật trong khoảng". Về mặt triết học, cách '
    'hiểu này SAI cho CI nhưng ĐÚNG cho CrI. Tuy nhiên trong thực hành, cách hiểu "loose" này '
    'ít khi dẫn đến quyết định lâm sàng sai.',
    color_hex='FFF3E0')
P()

# ============================================================
# PHẦN 3 — ĐỌC MỘT KẾT QUẢ CỤ THỂ TỪNG BƯỚC
# ============================================================
H2('3. Đọc một kết quả cụ thể — từng bước')

P('Lấy một kết quả thực từ báo cáo 10/04 của chúng ta:', italic=True)
P()

note_box(
    'KẾT QUẢ GỐC (Zhang et al. 2026, Bài 56, Bayesian meta-analysis 31 RCT, n = 19.865 HS):\n\n'
    '"School-based CBT produced a small reduction in anxiety symptoms\n'
    'SMD: –0.19, 95 % CrI: –0.22 to –0.17"',
    color_hex='F5F5F5')
P()

H3('BƯỚC 1 — Hiểu ký hiệu SMD là gì')
PRun([
    ('SMD = ', {'bold': True}),
    ('S', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('tandardized ', {}),
    ('M', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('ean ', {}),
    ('D', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('ifference = Khác biệt trung bình CHUẨN HOÁ', {}),
])
P()
P('Tại sao phải CHUẨN HOÁ? Vì 31 nghiên cứu được tổng hợp đã dùng NHIỀU thang đo lo âu khác '
  'nhau: GAD-7 (thang 0–21), SCARED (0–82), RCADS (0–45), STAI-C (20–80)... Không thể cộng '
  'trực tiếp điểm trung bình từ các thang này — phải chia cho độ lệch chuẩn (SD) để quy về '
  'CÙNG ĐƠN VỊ (đơn vị SD).')
P()
PRun([
    ('Công thức đơn giản: ', {'italic': True}),
    ('SMD = (Trung bình nhóm CT – Trung bình nhóm ĐC) ÷ SD chung', {'bold': True, 'italic': True}),
])
P()

H3('BƯỚC 2 — Hiểu con số –0,19')
P('Giá trị SMD = –0,19 nghĩa là:', italic=True)
P('• Nhóm CAN THIỆP (nhận CBT) có điểm lo âu trung bình THẤP HƠN nhóm ĐỐI CHỨNG một khoảng '
  'bằng 0,19 × SD.')
P('• Dấu âm (–) = can thiệp LÀM GIẢM triệu chứng (so với đối chứng tăng hoặc giữ nguyên).')
P('• Nếu thầy đổi sang ngôn ngữ đời thường: hiệu quả của CBT tương đương ≈ 20 % của một độ '
  'lệch chuẩn — rất NHỎ.')
P()
P('Áp dụng thang Cohen để đánh giá độ lớn:', bold=True)
MakeTable(
    ['Độ lớn |SMD|', 'Mức độ', 'Ví dụ thực tế'],
    [
        ('< 0,2', 'TRIVIAL (không đáng kể)',
         'Như việc một người hôm qua cao 170,0 cm → hôm nay đo 170,1 cm'),
        ('0,2 – 0,5', 'NHỎ (small)',
         'Khác biệt giữa HS giỏi vs HS khá môn Toán (đủ nhận ra nhưng không lớn)'),
        ('0,5 – 0,8', 'TRUNG BÌNH (medium)',
         'Khác biệt chiều cao giữa nam và nữ 15 tuổi (nhìn là thấy)'),
        ('> 0,8', 'LỚN (large)',
         'Khác biệt chiều cao giữa người lớn và trẻ 10 tuổi (khác biệt rõ ràng)'),
    ])
P()
PRun([
    ('→ SMD = –0,19 ', {'bold': True}),
    ('RẤT GẦN NGƯỠNG TRIVIAL', {'bold': True, 'red': True}),
    (', gần như không đáng kể về lâm sàng.', {}),
])
P()

H3('BƯỚC 3 — Hiểu khoảng [–0,22; –0,17]')
P('Đây là "KTC 95 % CrI" (Credible Interval 95 %).', italic=True)
P()
P('Hiểu thế này:', bold=True)
P('"Ước lượng điểm tốt nhất của hiệu quả là –0,19. Nhưng vì có thể có sai số, giá trị THỰC '
  'có thể là –0,18, có thể là –0,21, có thể bất cứ đâu quanh –0,19. Chúng ta không chắc chắn '
  '100 %, nhưng có thể nói: CÓ 95 % XÁC SUẤT giá trị thực nằm giữa –0,22 và –0,17."')
P()
PRun([
    ('Nếu thay vào tuyên bố lâm sàng: ', {'bold': True}),
    ('"Effect có thể nằm bất cứ đâu từ –0,22 (nhỏ) đến –0,17 (rất nhỏ) — nhưng chúng ta 95 % '
     'chắc rằng effect KHÔNG LỚN HƠN –0,22 cũng KHÔNG NHỎ HƠN –0,17."', {'italic': True}),
])
P()

H3('BƯỚC 4 — Hai câu hỏi QUAN TRỌNG cho thầy khi đọc')

P('Câu hỏi 1: Khoảng [–0,22; –0,17] có CHỨA SỐ 0 không?', bold=True)
P('→ KHÔNG. Cả hai đầu đều âm. Nghĩa là: hiệu quả THỰC SỰ khác 0, có ý nghĩa thống kê.',
  color=RGBColor(0x0B, 0x5E, 0x20))
P()
P('Câu hỏi 2: Độ lớn effect có ĐÁNG KỂ về lâm sàng không?', bold=True)
PRun([
    ('→ ', {}),
    ('KHÔNG ĐÁNG KỂ', {'bold': True, 'red': True}),
    ('. |SMD| từ 0,17 đến 0,22 ĐỀU dưới ngưỡng 0,2 của Cohen → TRIVIAL / NHỎ MỜ NHẠT. '
     'Dù có ý nghĩa thống kê, KHÔNG có ý nghĩa thực tiễn lớn.', {'red': True}),
])
P()

note_box(
    'KẾT LUẬN LÂM SÀNG: "CBT universal (phổ quát, cho mọi HS không chọn lọc) trong trường học '
    'cho thấy hiệu quả giảm lo âu có ý nghĩa THỐNG KÊ nhưng KHÔNG đủ lớn về LÂM SÀNG để khuyến '
    'nghị triển khai rộng rãi. Nên chọn hình thức TARGETED (chỉ cho HS có triệu chứng lo âu) '
    'để tối đa hoá hiệu quả."\n\n'
    'Đây chính xác là kết luận của tác giả Zhang 2026 — và việc thầy đọc được KTC 95 % CrI '
    'cho phép thầy TỰ CHẨN ĐOÁN lập luận này, không cần tin 100 % vào kết luận tác giả.',
    color_hex='E8F5E9')
P()

# ============================================================
# PHẦN 4 — KHI NÀO CrI KHÁC BIỆT VỚI CI
# ============================================================
H2('4. Khi nào CrI thực sự khác biệt với CI? (không phải lúc nào cũng khác)')

P('Một sự thật ít được nói trong sách giáo khoa: trong 90 % trường hợp meta-analysis tâm lý '
  'học hiện đại, SỐ CỦA CI VÀ CrI TRÙNG NHAU đến 2–3 chữ số thập phân. Vì sao?', italic=True)
P()

H3('Điều kiện khi KTC CI và CrI gần như TRÙNG NHAU:')
P('• Cỡ mẫu lớn (n > 100, đặc biệt tổng n meta-analysis > 500)', size=11)
P('• Prior của Bayesian là "non-informative" hoặc "weakly informative" (mặc định trong đa số phần mềm)', size=11)
P('• Phân phối dữ liệu gần chuẩn (không lệch cực trái/phải)', size=11)
PRun([
    ('→ Đây là điều kiện MẶC ĐỊNH của đa số psych MA hiện nay. Do đó đọc "CrI" ≈ đọc "CI" về SỐ.',
     {'bold': True, 'italic': True})
])
P()

H3('Điều kiện khi CrI THỰC SỰ KHÁC CI:')
P('• Cỡ mẫu nhỏ (n < 30)', size=11)
P('• Prior informative mạnh (ví dụ: nhà NC đưa niềm tin ban đầu rất chắc chắn từ NC trước)', size=11)
P('• Phân phối dữ liệu không chuẩn (lệch nặng, skewed)', size=11)
PRun([
    ('→ Trong các trường hợp này, CrI trả lời ', {}),
    ('chính xác hơn', {'bold': True}),
    (' và có thể khác CI rõ rệt.', {}),
])
P()

note_box(
    'TRONG THƯ VIỆN DỰ ÁN LO ÂU: 3 bài Bayesian (QT049 Zhang 2026, QT029 Li 2025, QT039 Xian '
    '2024) đều có tổng n > 500 (19.865, 1.711, 1.547) và dùng prior "weakly informative". Do '
    'đó CrI trong 3 bài này có thể được đọc "như CI thông thường" mà KHÔNG sai về diễn giải '
    'lâm sàng. Điểm KHÁC BIỆT ĐỘC ĐÁO của Bayesian NMA (QT029, QT039) là cho phép tính SUCRA '
    '— xác suất xếp hạng — dễ truyền thông hơn cho chính sách.',
    color_hex='E1F5FE')
P()

# ============================================================
# PHẦN 5 — LÀM THẾ NÀO KHI THẦY VIẾT BÁO CÁO
# ============================================================
H2('5. Khi thầy tự viết báo cáo — nên trích như thế nào?')

H3('Khi trích kết quả Bayesian (có CrI):')
P('Cách 1 — Ghi đầy đủ (tạp chí học thuật quốc tế):', italic=True)
PRun([
    ('✅ "SMD = –0,19 (95 % CrI: –0,22 đến –0,17)"', {'color': RGBColor(0x0B, 0x5E, 0x20), 'italic': True}),
])
P()
P('Cách 2 — Tiếng Việt (báo cáo cho đồng nghiệp VN):', italic=True)
PRun([
    ('✅ "SMD = –0,19 (KTC Bayesian 95 %: –0,22 đến –0,17)"', {'color': RGBColor(0x0B, 0x5E, 0x20), 'italic': True}),
])
P()
P('Cách 3 — Ngắn gọn trong bảng:', italic=True)
PRun([
    ('✅ "–0,19 [–0,22; –0,17]" với chú thích "Giá trị trong ngoặc là 95 % CrI (khoảng tin cậy Bayesian)".',
     {'color': RGBColor(0x0B, 0x5E, 0x20), 'italic': True}),
])
P()

H3('Nên tránh:')
PRun([
    ('❌ "SMD = –0,19 (KTC 95 %: –0,22; –0,17)" — mơ hồ, không rõ CI hay CrI',
     {'color': RGBColor(0xC0, 0x00, 0x00), 'italic': True}),
])
PRun([
    ('❌ "SMD = –0,19, p < 0,05" — thiếu khoảng tin cậy (APA 7 yêu cầu)',
     {'color': RGBColor(0xC0, 0x00, 0x00), 'italic': True}),
])
P()

# ============================================================
# PHẦN 6 — TOP 5 KÝ HIỆU THƯỜNG GẶP
# ============================================================
H2('6. Top 5 ký hiệu hay gặp trong báo cáo tâm lý học')

MakeTable(
    ['Ký hiệu', 'Tên tiếng Việt', 'Khi thấy trong bài báo, hiểu thế này'],
    [
        ('M hoặc x̄', 'Trung bình',
         'Giá trị trung bình của mẫu. VD: M = 15,2 nghĩa điểm stress trung bình = 15,2'),
        ('SD',
         'Độ lệch chuẩn',
         'Đo độ PHÂN TÁN của dữ liệu. SD nhỏ = tụ gần trung bình; SD lớn = phân tán rộng'),
        ('CI / CrI 95 %',
         'Khoảng tin cậy 95 %',
         'Khoảng ước lượng của giá trị thật. Nếu KHÔNG chứa 0 → có ý nghĩa thống kê'),
        ('Cohen d / Hedges g / SMD',
         'Hiệu ứng chuẩn hoá',
         'Đo ĐỘ LỚN khác biệt giữa 2 nhóm, không phụ thuộc thang đo. Áp tiêu chí Cohen'),
        ('p',
         'p-value',
         'Nếu p < 0,05 → có ý nghĩa thống kê. Nhưng CHÚ Ý: p không cho biết ĐỘ LỚN effect'),
    ])
P()

note_box(
    'LƯU Ý CỰC QUAN TRỌNG: Một kết quả có p < 0,001 (rất có ý nghĩa thống kê) vẫn có thể có '
    'SMD chỉ 0,05 (trivial về lâm sàng). Ngược lại, SMD = 1,5 (hiệu quả rất lớn) vẫn có thể '
    'có p = 0,10 nếu n quá nhỏ. DO ĐÓ khi đọc báo cáo: luôn nhìn CẢ p-value VÀ effect size '
    '(hoặc CI/CrI). Chỉ nhìn p → dễ hiểu sai.',
    color_hex='FFEBEE')
P()

# ============================================================
# PHẦN 7 — CHECKLIST 5 CÂU HỎI
# ============================================================
H2('7. Checklist 5 câu hỏi khi thầy đọc báo cáo có KTC')

P('Khi gặp bất kỳ kết quả nào có dạng "estimate + KTC", hãy hỏi 5 câu:', italic=True)
P()

P('[ ] 1. Estimate là gì?', bold=True, size=12)
P('   → MD (raw)? SMD (chuẩn hoá)? OR? HR? β? r?', size=11, color=RGBColor(0x66, 0x66, 0x66))

P('[ ] 2. KTC là CI hay CrI?', bold=True, size=12)
P('   → Đọc Methods. Nếu thấy "Bayesian" → CrI. Nếu không → mặc định CI.', size=11, color=RGBColor(0x66, 0x66, 0x66))

P('[ ] 3. Khoảng có CHỨA 0 (hoặc 1 cho OR/HR) không?', bold=True, size=12)
P('   → Chứa 0 → không có ý nghĩa thống kê. Không chứa → có ý nghĩa.', size=11, color=RGBColor(0x66, 0x66, 0x66))

P('[ ] 4. Độ lớn effect có Ý NGHĨA LÂM SÀNG không?', bold=True, size=12)
P('   → Áp tiêu chí Cohen cho SMD, hoặc hiểu mối quan hệ với practice.', size=11, color=RGBColor(0x66, 0x66, 0x66))

P('[ ] 5. Heterogeneity (I², τ²) có được báo không?', bold=True, size=12)
P('   → Với MA: I² < 25 % = thấp (OK pool); 25–75 % = vừa; > 75 % = cao (pooled có thể misleading).',
  size=11, color=RGBColor(0x66, 0x66, 0x66))
P()

# ============================================================
# PHẦN 8 — TÓM TẮT CUỐI
# ============================================================
H2('8. Tóm tắt cuối cùng — nhớ 3 điều này là đủ')

P('1. ', bold=True, size=13)
PRun([
    ('KTC = Khoảng Tin Cậy. ', {'bold': True, 'size': 13}),
    ('Tiếng Anh có 2 kiểu: CI (tần suất) và CrI (Bayesian). Khi báo cáo dùng Bayesian MA → ghi CrI.', {'size': 12}),
])
P()

P('2. ', bold=True, size=13)
PRun([
    ('Trong thực hành lâm sàng, đọc CI và CrI GIỐNG NHAU. ', {'bold': True, 'size': 13}),
    ('Cả hai đều trả lời câu hỏi: "Giá trị thật của hiệu ứng có thể nằm đâu?". Kết luận lâm sàng '
     'hầu như không khác nhau.', {'size': 12}),
])
P()

P('3. ', bold=True, size=13)
PRun([
    ('Cái QUAN TRỌNG không phải CI vs CrI, mà là:', {'bold': True, 'size': 13}),
])
P('   (a) Khoảng có chứa 0 không? → ý nghĩa thống kê', size=12)
P('   (b) Effect size có đủ lớn không? → ý nghĩa lâm sàng (dùng Cohen cutoffs)', size=12)
P('   (c) Heterogeneity có cao không? → pooled estimate đáng tin không', size=12)
P()

note_box(
    'NGẮN GỌN NHẤT: Khi đọc "SMD = –0,19 (95 % CrI: –0,22; –0,17)", thầy có thể hiểu đơn giản '
    'như: "Ước lượng tốt nhất của hiệu quả là –0,19 (hiệu ứng nhỏ, âm → can thiệp có giảm '
    'triệu chứng), với độ tin cậy 95 % giá trị thực nằm trong khoảng –0,22 đến –0,17 (cũng '
    'đều nhỏ). Khoảng không chứa 0 nên hiệu quả có thật, nhưng quá nhỏ nên không đáng triển '
    'khai universal — ưu tiên targeted."',
    color_hex='E8F5E9')
P()

# ============================================================
# PHỤ LỤC
# ============================================================
P('─' * 70, color=RGBColor(0x99, 0x99, 0x99))
H3('Phụ lục: Các bài Bayesian trong thư viện dự án Lo âu')

MakeTable(
    ['ID', 'Tác giả / Năm', 'Loại nghiên cứu', 'Ví dụ kết quả có CrI'],
    [
        ('QT049', 'Zhang, Liang & Kang 2026',
         'Bayesian MA 31 RCT CBT học đường',
         'SMD lo âu –0,19 (95 % CrI: –0,22; –0,17); duy trì đến 1 năm'),
        ('QT029', 'Li et al. 2025',
         'Bayesian NMA 30 RCT điều trị lo âu trẻ em',
         'CBT xếp hạng 2 (SUCRA 0,66) trong 9 liệu pháp'),
        ('QT039', 'Xian, Zhang & Jiang 2024',
         'Bayesian NMA 30 RCT SAD ở VTN',
         'iCBT xếp hạng 1 (SUCRA 71,2 %)'),
    ])
P()

P('Tài liệu v3 — ngày 15/04/2026. Tập trung vào sự dễ hiểu cho người thực hành lâm sàng/'
  'giáo viên tâm lý học, không phải thống kê gia.',
  italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT):,} bytes')
