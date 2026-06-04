# -*- coding: utf-8 -*-
"""Dịch đầy đủ QT27 — Nature Human Behaviour 2025 — Fassi et al. — Social Media & MH"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link bài báo gốc: https://doi.org/10.1038/s41562-025-02134-4', size=10)

add_heading(doc, 'Sử dụng mạng xã hội ở thanh thiếu niên có và không có tình trạng sức khỏe tâm thần', 1)
h = doc.add_paragraph()
r = h.add_run('Social Media Use in Adolescents with and without Mental Health Conditions')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Social Media Use in Adolescents with and without Mental Health Conditions'),
    ('Tiêu đề dịch', 'Sử dụng mạng xã hội ở thanh thiếu niên có và không có tình trạng sức khỏe tâm thần'),
    ('Tác giả', 'Luisa Fassi, Amanda M. Ferguson, Andrew K. Przybylski, Tamsin J. Ford, Amy Orben'),
    ('Cơ quan', 'MRC Cognition and Brain Sciences Unit, University of Cambridge; Department of Psychiatry, Cambridge; Oxford Internet Institute, University of Oxford'),
    ('Tạp chí', 'Nature Human Behaviour (Q1, IF ≈ 30)'),
    ('Thông tin xuất bản', '2025, Vol. 9, pp. 1283–1299, 21 trang'),
    ('DOI', '10.1038/s41562-025-02134-4'),
    ('Loại NC', 'Registered Report — phân tích cắt ngang dữ liệu MHCYP 2017 (NHS Digital)'),
    ('Mẫu', 'N = 3.340 VTN Anh (11–19 tuổi; TB = 14,77; SD = 2,48), đại diện quốc gia'),
    ('Nhận bài', '23/12/2022; Chấp nhận: 11/02/2025; Xuất bản: 05/05/2025'),
])
add_page_ref(doc, '1283–1299', 'Nature Human Behaviour', 'Vol. 9, June 2025')

# TÓM TẮT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Lo ngại về mối quan hệ giữa sử dụng MXH và SKTT VTN đang tăng, nhưng ít NC tập trung vào VTN có triệu chứng SKTT ở mức lâm sàng. Điều này hạn chế hiểu biết về cách sử dụng MXH khác nhau giữa các hồ sơ SKTT khác nhau.')

p = doc.add_paragraph()
r = p.add_run('Trong Registered Report này, chúng tôi phân tích dữ liệu đại diện quốc gia UK (N = 3.340, 11–19 tuổi) bao gồm đánh giá chẩn đoán bởi nhà lâm sàng cùng với đo lường MXH định lượng và định tính. Như giả thuyết, VTN có tình trạng SKTT báo cáo dành NHIỀU thời gian hơn trên MXH và ÍT hạnh phúc hơn về số lượng bạn trực tuyến so với VTN không có tình trạng. Chúng tôi cũng tìm thấy khác biệt theo loại tình trạng: VTN có rối loạn NỘI HÓA báo cáo dành nhiều thời gian hơn, tham gia nhiều hơn vào so sánh xã hội, trải qua tác động lớn hơn của phản hồi lên tâm trạng, cùng với ít hạnh phúc hơn về số bạn trực tuyến và ít tự bộc lộ trung thực hơn. Ngược lại, VTN có rối loạn NGOẠI HÓA chỉ báo cáo thời gian cao hơn.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_p(doc, 'Những phát hiện này nhấn mạnh nhu cầu xem xét các hồ sơ SKTT đa dạng của VTN trong chính sách và thực hành lâm sàng.')

# ĐÁNH GIÁ NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Nature Human Behaviour Q1 IF ≈ 30 — tạp chí uy tín CỰC CAO trong khoa học hành vi.',
    'Registered Report — phương pháp TIỀN ĐĂNG KÝ: giả thuyết, phân tích được đánh giá và chấp nhận TRƯỚC KHI thu thập dữ liệu. Giảm thiên lệch xuất bản và p-hacking.',
    'Dữ liệu MHCYP (NHS Digital) — chẩn đoán LÂM SÀNG bởi nhà đánh giá (DAWBA), KHÔNG chỉ sàng lọc tự báo cáo. Đại diện quốc gia UK.',
    'Phân biệt NỘI HÓA vs NGOẠI HÓA — đóng góp lý thuyết quan trọng. VTN lo âu/trầm cảm (nội hóa) nhạy cảm MXH hơn VTN ADHD/hành vi (ngoại hóa).',
    'Đo cả ĐỊNH LƯỢNG (thời gian MXH) VÀ ĐỊNH TÍNH (so sánh xã hội, kiểm soát, phản hồi, bắt nạt mạng, bộc lộ bản thân) — toàn diện.',
    'Equivalence testing (TOST) — phương pháp tiên tiến. Không chỉ hỏi "có ý nghĩa thống kê?" mà còn hỏi "khác biệt có ĐỦ LỚN để quan trọng?".',
    'N = 3.340, 50% nam/50% nữ, 16% có ≥1 tình trạng SKTT (chẩn đoán).',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Cắt ngang — KHÔNG xác lập nhân quả (MXH → SKTT hay SKTT → MXH?). Li 2025 (QT22) dùng dọc nhưng cũng yếu.',
    'Chỉ UK 2017 — TRƯỚC TikTok bùng nổ (2019+). Bối cảnh MXH 2025 rất khác 2017.',
    '16% có SKTT — mẫu lâm sàng nhỏ (n = 519 tổng; nội hóa n = 282; ngoại hóa n = 104).',
    'Tự báo cáo thời gian MXH — không đo khách quan. JAMA 2024 (QT33) dùng tracker khách quan.',
    'Chỉ MXH — không đo game, TV, screen time tổng. Hoàng Trung Học 2025 (VN14): game β = 0,176 riêng biệt.',
]:
    add_p(doc, f'• {b}')

# 1. GIỚI THIỆU
add_page_ref(doc, '1283–1284', 'Nature Human Behaviour', '2025')
add_heading(doc, '1. GIỚI THIỆU', 2)

add_p(doc, 'VTN trên toàn thế giới đã trải qua suy giảm SKTT trong thập kỷ qua. Dữ liệu UK gần đây cho thấy 1/6 trẻ 7–16 tuổi và 1/4 trẻ 17–19 tuổi có tình trạng SKTT có thể — tăng rõ rệt từ 1/9 và 1/10 năm 2017. Vì 48% người có tình trạng SKTT lần đầu trải nghiệm triệu chứng trước 18 tuổi (Solmi et al. 2022), gánh nặng SKTT tăng này sẽ ảnh hưởng tiêu cực đến xã hội, kinh tế, và cuộc sống VTN + người trưởng thành. Nhiều người lo ngại xu hướng này do MXH gây ra — 93% VTN 12–17 tuổi hiện có tài khoản MXH.')

add_p(doc, 'Tuy nhiên, nhiều nhà nghiên cứu đặt câu hỏi về sức mạnh bằng chứng hiện tại. Tài liệu cho nhiều kết quả mâu thuẫn. Tranh luận về: thiếu bằng chứng dọc/nhân quả, kích thước hiệu ứng nào quan trọng (Orben & Przybylski, 2019 vs Twenge & Haidt, 2020), và sự khác biệt cá nhân lớn theo tuổi, giới, dân tộc.')

add_p(doc, 'Hạn chế chính: hầu hết NC dùng MẪU KHÔNG LÂM SÀNG — VTN được tuyển từ trường học hoặc cộng đồng, đánh giá SKTT bằng bảng hỏi sàng lọc. Rất ít NC bao gồm VTN có triệu chứng ở mức LÂM SÀNG (tức là đạt tiêu chuẩn chẩn đoán rối loạn). Điều này nghiêm trọng vì: (1) VTN có rối loạn có thể sử dụng và bị ảnh hưởng bởi MXH KHÁC, (2) chính sách và can thiệp cần dựa trên bằng chứng từ đúng nhóm đích.')

add_p(doc, 'Ngoài ra, rối loạn tâm thần rất ĐA DẠNG. Rối loạn nội hóa (internalizing: lo âu, trầm cảm) biểu hiện tiêu cực hướng NỘI TÂM — suy nghĩ lặp đi lặp lại, lo lắng, rút lui xã hội. Rối loạn ngoại hóa (externalizing: ADHD, rối loạn hành vi) biểu hiện tiêu cực hướng NGƯỜI KHÁC — bốc đồng, mạo hiểm, thiếu ức chế. NC sử dụng bảng hỏi chọn lọc không thể bao quát đầy đủ sự đa dạng lâm sàng này.')

add_p(doc, 'Registered Report này cung cấp dữ liệu và bằng chứng quan trọng về mối quan hệ MXH và SKTT giữa các nhóm VTN đáp ứng và không đáp ứng tiêu chuẩn chẩn đoán cho nhiều loại tình trạng SKTT. Sử dụng dữ liệu MHCYP — VTN KHÔNG được tuyển hoặc chẩn đoán từ phòng khám, mà trải qua đánh giá SKTT như phần của khảo sát MHCYP.')

# 2. PHƯƠNG PHÁP
add_page_ref(doc, '1285–1288', 'Nature Human Behaviour', '2025')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)

add_p(doc, '2.1. Dữ liệu và mẫu', bold=True)
add_p(doc, 'MHCYP 2017 (Mental Health of Children and Young People in England) — khảo sát quốc gia do NHS Digital thực hiện. N = 3.340 VTN 11–19 tuổi (TB = 14,77; SD = 2,48). 50% nam, 50% nữ. Đại diện quốc gia.')
add_p(doc, 'Chẩn đoán SKTT: DAWBA (Development and Well-Being Assessment) — đánh giá bởi nhà lâm sàng được đào tạo, dựa trên thông tin từ VTN, cha mẹ, và giáo viên. KHÔNG phải tự báo cáo.')
add_p(doc, '16% có ≥1 tình trạng SKTT (N = 519). Phân loại: Nội hóa (lo âu, trầm cảm, OCD, eating disorders; N = 282, 8%); Ngoại hóa (ADHD, rối loạn hành vi; N = 104, 3%). Một số có cả hai.')

add_p(doc, '2.2. Đo lường MXH', bold=True)
add_p(doc, '7 chiều MXH được đo (Bảng 2 gốc), dựa trên y văn:')
add_p(doc, '• (H1.1a) Thời gian trên MXH: "Khi bạn dùng MXH, tổng thời gian bao lâu trong ngày học/cuối tuần?" — Đo định lượng.')
add_p(doc, '• (H1.1b) So sánh xã hội trực tuyến: "Bạn có so sánh bản thân với người khác trên MXH?"')
add_p(doc, '• (H1.1c) Thiếu kiểm soát thời gian: "Bạn có thấy khó ngừng dùng MXH?"')
add_p(doc, '• (H1.1d) Bộc lộ bản thân trung thực: "Bạn có thể là chính mình trên MXH?"')
add_p(doc, '• (H1.1e) Tác động phản hồi lên tâm trạng: "Số likes, comments, shares có ảnh hưởng tâm trạng bạn?"')
add_p(doc, '• (H1.2a) Hạnh phúc về số bạn trực tuyến.')
add_p(doc, '• (H1.2b) Bắt nạt mạng.')

add_p(doc, '2.3. Phân tích', bold=True)
add_p(doc, 'Registered Report — giả thuyết và phân tích TIỀN ĐĂNG KÝ và được chấp nhận TRƯỚC khi truy cập dữ liệu. Giảm thiên lệch xuất bản và p-hacking.')
add_p(doc, 'So sánh nhóm: Hedges\' g (kích thước hiệu ứng). NHST (p-value) + Equivalence Testing (TOST — Two One-Sided Tests). SESOI (Smallest Effect Size of Interest) = g = 0,4 — ngưỡng nhỏ nhất đáng quan tâm.')
add_p(doc, 'Nếu g > 0,4 VÀ TOST không tương đương: khác biệt CÓ Ý NGHĨA THỰC TẾ.')
add_p(doc, 'Nếu g < 0,4 VÀ TOST tương đương: khác biệt CÓ nhưng NHỎ (dưới ngưỡng quan trọng).')
add_p(doc, 'Điều chỉnh cho tuổi và giới. Hồi quy tuyến tính với clustering.')

# 3. KẾT QUẢ
add_page_ref(doc, '1288–1293', 'Nature Human Behaviour', '2025')
add_heading(doc, '3. KẾT QUẢ', 2)

add_p(doc, '3.1. Mô tả mẫu', bold=True)
add_p(doc, 'N = 3.340 VTN 11–19 tuổi (TB 14,77; SD 2,48). 50% nam, 50% nữ. 519 (16%) có ≥1 tình trạng SKTT. Nội hóa: 282 (8%). Ngoại hóa: 104 (3%).')

add_p(doc, '3.2. BẤT KỲ tình trạng SKTT vs KHÔNG (H1)', bold=True)

add_heading(doc, 'Bảng 1. So sánh sử dụng MXH: VTN có vs không có tình trạng SKTT', 3)
add_table(doc,
    ['Chiều MXH', 'Hedges\' g', 'KTC 90%', 'NHST β (SE)', 'p', 'Tương đương?', 'Kết luận'],
    [['(a) Thời gian MXH', '0,46', '0,38–0,54', '0,87 (0,08)', '2,7×10⁻²⁵', 'KHÔNG (>SESOI)', 'KHÁC BIỆT CÓ Ý NGHĨA'],
     ['(b) So sánh xã hội', '0,30', '0,22–0,39', '0,42 (0,07)', '8,0×10⁻¹¹', 'Có (<SESOI)', 'Có nhưng NHỎ'],
     ['(c) Thiếu kiểm soát', '0,27', '0,19–0,35', '0,39 (0,07)', '3,6×10⁻⁸', 'Có', 'Có nhưng NHỎ'],
     ['(d) Bộc lộ trung thực', '−0,12', '−0,20–−0,04', '—', '—', 'Có', 'NHỎ (hướng ngược)'],
     ['(e) Tác động phản hồi', '0,29', '0,21–0,38', '0,38 (0,06)', '2,3×10⁻¹⁰', 'Có', 'Có nhưng NHỎ'],
     ['Hạnh phúc bạn online', '−0,28', '−0,37–−0,20', '—', '—', 'Có', 'VTN SKTT ÍT hạnh phúc'],
     ['Bắt nạt mạng', '0,20', '0,12–0,28', '—', '—', 'Có', 'NHỎ']],
    widths=[3.0, 1.5, 2.0, 2.0, 1.5, 2.0, 3.0])
add_p(doc, 'SESOI = g = 0,4. "Tương đương" = khác biệt < ngưỡng thực tế. CHỈ thời gian MXH (g = 0,46) vượt SESOI — tất cả chiều khác NHỎ hơn ngưỡng.', size=9, italic=True)

add_p(doc, '3.3. NỘI HÓA vs NGOẠI HÓA (H2, H3)', bold=True)

add_heading(doc, 'Bảng 2. So sánh nội hóa vs ngoại hóa', 3)
add_table(doc,
    ['Chiều MXH', 'Nội hóa g', 'Ngoại hóa g', 'Khác biệt', 'Ý nghĩa'],
    [['Thời gian MXH', '0,47', '0,38', 'Nội > ngoại', 'Cả hai cao hơn không SKTT'],
     ['So sánh xã hội', '0,42', '0,12', 'NỘI >> ngoại', 'Nội hóa nhạy cảm SO SÁNH'],
     ['Thiếu kiểm soát', '0,35', '0,26', 'Nội > ngoại', ''],
     ['Tác động phản hồi', '0,43', '0,08', 'NỘI >> ngoại', 'Nội hóa nhạy cảm PHẢN HỒI'],
     ['Bộc lộ trung thực', '−0,22', '0,04', 'NỘI thấp hơn', 'Nội hóa KHÔNG thật trên MXH'],
     ['Hạnh phúc bạn online', '−0,41', '−0,05', 'NỘI >> ngoại', 'Nội hóa RẤT bất hạnh'],
     ['Bắt nạt mạng', '0,24', '0,21', 'Tương đương', 'Cả hai bị bắt nạt']],
    widths=[3.0, 2.0, 2.0, 2.5, 4.0])
add_p(doc, 'Phát hiện THEN CHỐT: VTN NỘI HÓA (lo âu, trầm cảm) bị ảnh hưởng bởi MXH NHIỀU HƠN RÕ RỆT so với VTN NGOẠI HÓA (ADHD, hành vi). Ngoại hóa chỉ dùng MXH NHIỀU hơn (thời gian) nhưng không bị ảnh hưởng bởi NỘI DUNG.', size=9, italic=True)

add_p(doc, '3.4. Phân tích bổ sung theo loại rối loạn cụ thể', bold=True)
add_p(doc, 'Lo âu riêng: g = 0,50 cho thời gian MXH (cao nhất); 0,44 cho so sánh xã hội; 0,45 cho phản hồi. Lo âu NHẠY CẢM NHẤT với MXH trong tất cả rối loạn.')
add_p(doc, 'Trầm cảm: g = 0,38 cho thời gian; 0,40 cho so sánh; 0,39 cho phản hồi. Cũng nhạy cảm nhưng thấp hơn lo âu.')
add_p(doc, 'ADHD: g = 0,39 cho thời gian; 0,10 cho so sánh; 0,05 cho phản hồi. Dùng nhiều nhưng KHÔNG nhạy cảm nội dung.')
add_p(doc, 'Rối loạn hành vi: g = 0,34 cho thời gian; 0,14 cho so sánh. Tương tự ADHD.')

# 4. THẢO LUẬN
add_page_ref(doc, '1293–1297', 'Nature Human Behaviour', '2025')
add_heading(doc, '4. THẢO LUẬN', 2)

add_p(doc, 'Phát hiện chính: VTN có tình trạng SKTT dùng MXH nhiều hơn (g = 0,46) — khác biệt duy nhất vượt SESOI. Tất cả chiều MXH khác (so sánh, kiểm soát, phản hồi, bắt nạt) — có ý nghĩa thống kê nhưng NẰM TRONG NGƯỠNG TƯƠNG ĐƯƠNG. Nghĩa là: khác biệt TỒN TẠI nhưng CÓ THỂ KHÔNG ĐỦ LỚN ĐỂ QUAN TRỌNG về mặt thực tế.')

add_p(doc, 'Nội hóa vs ngoại hóa — đóng góp quan trọng nhất:', bold=True)
add_p(doc, 'VTN nội hóa (lo âu, trầm cảm) bị ảnh hưởng bởi MXH trên NHIỀU CHIỀU — thời gian, so sánh xã hội, phản hồi, bộc lộ, hạnh phúc bạn bè. VTN ngoại hóa (ADHD, hành vi) chỉ dùng MXH nhiều thời gian nhưng KHÔNG bị ảnh hưởng bởi nội dung. Giải thích: nội hóa liên quan đến suy nghĩ lặp đi lặp lại (rumination), lo lắng về đánh giá xã hội, thiên lệch nhận thức tiêu cực — tất cả được KHUẾCH ĐẠI bởi MXH (likes, comments, so sánh). Ngoại hóa liên quan đến bốc đồng, tìm kiếm kích thích — MXH cung cấp kích thích nhưng không gây rumination.')

add_p(doc, 'Hàm ý chính sách:', bold=True)
add_p(doc, 'Kết quả KHÔNG ủng hộ mạnh mẽ giả thuyết MXH "gây hại lớn" cho TẤT CẢ VTN (hầu hết khác biệt dưới SESOI). Tuy nhiên, VTN có rối loạn nội hóa (đặc biệt lo âu) là NHÓM NGUY CƠ CAO — cần can thiệp nhắm vào nhóm này, không phải hạn chế MXH phổ quát cho tất cả.')

add_p(doc, 'Liên hệ với các NC khác trong Đề tài:', bold=True)
add_p(doc, '• Norway 2025 (QT21): MXH giải thích MỘT PHẦN xu hướng tăng distress — phù hợp (đóng góp nhưng không phải duy nhất).')
add_p(doc, '• Li 2025 (QT22): Tác động dọc YẾU — phù hợp (hầu hết khác biệt NHỎ trong bài này).')
add_p(doc, '• JAMA 2024 (QT33): RCT giảm screen time → cải thiện nội hóa mạnh nhất (d = 0,53) — XÁC NHẬN: nội hóa nhạy cảm MXH/screen time nhất.')
add_p(doc, '• Hoàng Trung Học 2025 (VN14): Game điện tử β = 0,176 — chỉ đo thời gian, chưa đo nội dung MXH.')

add_p(doc, 'Hạn chế:', bold=True)
add_p(doc, '• Cắt ngang — KHÔNG xác lập nhân quả. MXH → SKTT hay SKTT → MXH (hoặc cả hai)? VTN lo âu có thể TĂNG dùng MXH vì rút lui xã hội trực tiếp (Wadley et al. 2020). Li 2025 (QT22) gợi ý quan hệ hai chiều.')
add_p(doc, '• Dữ liệu 2017 — TRƯỚC TikTok (ra mắt quốc tế 2018). Bối cảnh MXH 2025 rất khác: TikTok, Reels, short-form video. Các NC tương lai cần dữ liệu mới hơn.')
add_p(doc, '• 16% có SKTT — mẫu lâm sàng tương đối nhỏ. Nội hóa n = 282, ngoại hóa n = 104 — hạn chế phân tích tiểu nhóm sâu hơn (ví dụ: lo âu riêng n có thể <150).')
add_p(doc, '• Tự báo cáo thời gian MXH — không đo khách quan. JAMA 2024 (QT33) dùng tracker — phương pháp tốt hơn.')
add_p(doc, '• Chỉ đo MXH — không đo game, TV, screen time tổng. Hoàng Trung Học 2025 (VN14): game riêng biệt (β = 0,176). Cần phân biệt loại screen media.')
add_p(doc, '• Chỉ UK — văn hóa MXH có thể khác châu Á/VN (Zalo, TikTok VN vs Instagram UK).')

# 5. KẾT LUẬN
add_heading(doc, '5. KẾT LUẬN', 2)
add_p(doc, 'Dữ liệu 3.340 VTN Anh (11–19 tuổi) với chẩn đoán lâm sàng DAWBA cho thấy:')
add_p(doc, '(1) VTN có tình trạng SKTT dùng MXH nhiều hơn (g = 0,46) — khác biệt duy nhất vượt ngưỡng thực tế (SESOI 0,4). Các chiều MXH khác: có ý nghĩa thống kê nhưng NHỎ.')
add_p(doc, '(2) VTN NỘI HÓA (lo âu, trầm cảm) nhạy cảm MXH trên NHIỀU CHIỀU — so sánh xã hội (g = 0,42), phản hồi (g = 0,43), bộc lộ (g = −0,22). VTN NGOẠI HÓA chỉ dùng nhiều thời gian.')
add_p(doc, '(3) Kết quả KHÔNG ủng hộ giả thuyết MXH "gây hại lớn" cho tất cả VTN, nhưng xác nhận VTN lo âu là NHÓM NGUY CƠ CAO cần can thiệp nhắm mục tiêu.')
add_p(doc, 'Gợi ý rằng chính sách cần NHẮM MỤC TIÊU (VTN nội hóa) thay vì HẠN CHẾ PHỔ QUÁT, và đo lường MXH cần cả ĐỊNH LƯỢNG lẫn ĐỊNH TÍNH.')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
refs = [
    'Fassi, L., Ferguson, A.M., Przybylski, A.K., Ford, T.J. & Orben, A. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9, 1283–1299.',
    'Orben, A. & Przybylski, A.K. (2019). The association between adolescent well-being and digital technology use. Nature Human Behaviour, 3(2), 173–182.',
    'Twenge, J.M. & Haidt, J. (2020). Underestimating digital media harm. Nature Human Behaviour, 4(4), 346–348.',
    'Solmi, M., et al. (2022). Age at onset of mental disorders worldwide. Molecular Psychiatry, 27, 281–295.',
    'Orben, A., Meier, A., Dalgleish, T. & Blakemore, S.J. (2024). Mechanisms linking social media use to adolescent mental health vulnerability. Nature Reviews Psychology, 3, 407–423.',
    'Goodyer, I.M., Reynolds, S., Barrett, B., et al. (2017). Cognitive-behavioural therapy and short-term psychoanalytical psychotherapy versus a brief psychosocial intervention. Lancet Psychiatry, 4, 529–540.',
    'Li, S.H., et al. (2025). Cross-sectional and longitudinal associations of screen time with adolescent depression and anxiety. BJCP, 64, 873–887.',
    'Schmidt-Persson, J., et al. (2024). Screen media use and mental health of children and adolescents. JAMA Network Open, 7(7), e2419881.',
    'Brunborg, G.S., et al. (2025). Possible explanations for the upward trend in mental distress among adolescents in Norway. Social Science & Medicine, 384, 118528.',
]
for ref in refs:
    add_p(doc, ref, size=10)
add_p(doc, '(Xem danh mục đầy đủ trong bài gốc — 78 tài liệu tham khảo)', size=10, italic=True)

# VIẾT TẮT
add_abbreviation_table(doc, [
    ('MXH', 'Mạng xã hội (Social Media)'),
    ('SKTT', 'Sức khỏe tâm thần'),
    ('VTN', 'Vị thành niên'),
    ('MHCYP', 'Mental Health of Children and Young People in England'),
    ('DAWBA', 'Development and Well-Being Assessment — Đánh giá Phát triển và Hạnh phúc'),
    ('SDQ', 'Strengths and Difficulties Questionnaire'),
    ('SESOI', 'Smallest Effect Size of Interest — Kích thước Hiệu ứng Nhỏ nhất Đáng quan tâm'),
    ('NHST', 'Null Hypothesis Significance Testing — Kiểm định Ý nghĩa Giả thuyết Không'),
    ('TOST', 'Two One-Sided Tests — Kiểm định Tương đương Hai phía'),
    ('g', 'Hedges\' g — Kích thước hiệu ứng (tương tự Cohen d, điều chỉnh mẫu nhỏ)'),
    ('ADHD', 'Attention-Deficit/Hyperactivity Disorder — Rối loạn Tăng động Giảm chú ý'),
    ('OCD', 'Obsessive-Compulsive Disorder — Rối loạn Ám ảnh Cưỡng chế'),
    ('NHS', 'National Health Service — Dịch vụ Y tế Quốc gia Anh'),
])

# PHẢN BIỆN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'Nature Human Behaviour Q1 IF ≈ 30 — CỰC UY TÍN. Registered Report — tiêu chuẩn vàng cho NC không thiên lệch.',
    'Chẩn đoán LÂM SÀNG (DAWBA) — vượt trội so với TOÀN BỘ NC khác trong Đề tài (tất cả dùng sàng lọc: GAD-7, DASS-21, K10). V-NAMHS 2022 dùng DISC-5 — cũng chẩn đoán, nhưng DAWBA đánh giá bởi nhà lâm sàng con người.',
    'Equivalence testing (TOST) — phương pháp TIÊN TIẾN NHẤT. Phân biệt rõ "có ý nghĩa thống kê" (p < 0,05) vs "đủ lớn để quan trọng" (> SESOI). Ít NC nào dùng — 59 Countries (QT31), Norway (QT21), JAACAP (QT23) đều chỉ dùng p-value.',
    'Phân biệt NỘI HÓA vs NGOẠI HÓA — đóng góp lý thuyết ĐỘC ĐÁO. Xác nhận cơ chế: rumination + so sánh tiêu cực (nội hóa) bị khuếch đại bởi MXH. Phù hợp AJP 2024 (QT28, Zugman et al.): cơ chế lo âu liên quan thiên lệch chú ý + xử lý đe dọa.',
    'Đo CẢ ĐỊNH LƯỢNG (thời gian) VÀ ĐỊNH TÍNH (7 chiều MXH) — toàn diện nhất trong Đề tài. Li 2025 (QT22) chỉ đo 1 câu thời gian; Norway (QT21) chỉ đo thời gian; Hoàng Trung Học (VN14) chỉ đo game.',
    'JAMA 2024 (QT33) xác nhận: RCT giảm screen time → cải thiện NỘI HÓA mạnh nhất (−1,03). Bài này giải thích TẠI SAO: VTN nội hóa nhạy cảm hơn trên nhiều chiều MXH.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'CẮT NGANG — hạn chế lớn nhất. KHÔNG xác lập nhân quả. VTN lo âu có thể TĂNG dùng MXH (rút lui xã hội → online nhiều hơn) thay vì MXH GÂY lo âu. Li 2025 (QT22): tác động dọc YẾU (lo âu p = 0,443), gợi ý quan hệ hai chiều. Cần ecological momentary assessment (EMA) hoặc RCT.',
    'Dữ liệu 2017 — TRƯỚC TikTok bùng nổ (2019+). MXH 2025 rất khác: short-form video, algorithm cá nhân hóa, doom scrolling. Ireland 2024 (QT32) cũng lưu ý: thang MXH tiêu cực (Facebook 2013) lỗi thời. Cần lặp lại với dữ liệu 2024+.',
    '16% có SKTT (N = 519) — mẫu lâm sàng TỐT nhưng vẫn nhỏ cho phân tích tiểu nhóm sâu. Nội hóa n = 282 → phân tích lo âu riêng có thể thiếu sức mạnh. 59 Countries (QT31) n = 179.937 nhưng chỉ dùng 1 câu lo âu.',
    'Tự báo cáo thời gian MXH — thiên lệch hồi tưởng. JAMA 2024 (QT33) dùng tracker KHÁCH QUAN (app, phần mềm, TV monitor) — phương pháp tốt hơn.',
    'Chỉ UK — văn hóa MXH rất khác VN (Instagram/Snapchat UK vs Zalo/TikTok VN). Hoàng Trung Học 2025 (VN14): game (không phải MXH) có β = 0,176. VN cần NC riêng.',
    'SESOI g = 0,4 — ngưỡng do tác giả chọn trước, có thể tranh cãi. Một số cho rằng g = 0,2–0,3 đã quan trọng với dân số lớn (Funder & Ozer, 2019). Nếu dùng SESOI = 0,3: so sánh xã hội (0,30) VÀ phản hồi (0,29) sẽ ở NGƯỠNG.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống nghiên cứu / Research Gap:', bold=True)
for s in [
    'VN CHƯA CÓ NC MXH với chẩn đoán lâm sàng — tất cả dùng sàng lọc (GAD-7, DASS-21). Cần DAWBA/DISC cho VTN VN để phân biệt "có triệu chứng" vs "có rối loạn".',
    'Cần NC NỘI HÓA vs NGOẠI HÓA ở VN — VTN VN lo âu (nội hóa) có nhạy cảm MXH hơn ADHD (ngoại hóa) không? Phù hợp cho can thiệp nhắm mục tiêu.',
    'Đo cả ĐỊNH LƯỢNG VÀ ĐỊNH TÍNH MXH ở VN (không chỉ thời gian) — so sánh xã hội, phản hồi, bộc lộ. Bài này cho thấy thời gian ALONE không đủ để hiểu mối quan hệ.',
    'So sánh tác động MXH ở VN vs UK — văn hóa khác (Zalo, TikTok VN). Hoàng Trung Học 2025 (VN14): game β = 0,176 nhưng MXH chưa đo riêng. Jefferies 2020 (QT35): VN SAD = 30,7% — lo âu XH có thể liên quan MXH.',
    'Ecological momentary assessment (EMA) tại VN — đo MXH + tâm trạng THỜI GIAN THỰC, nhiều lần/ngày. Giải quyết vấn đề nhân quả mà cắt ngang và dọc đều gặp khó.',
    'RCT nhắm mục tiêu VTN NỘI HÓA: giảm so sánh xã hội + phản hồi (không chỉ giảm thời gian tổng). BMC NMA 2025 (QT29): ACT (SUCRA 0,69) có thể giúp VTN lo âu chấp nhận cảm xúc khi dùng MXH thay vì tránh hoàn toàn.',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '27_NatureHumanBehav_SocialMedia_2025.docx')
doc.save(outpath)
import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
