# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0,0,0)

def p(text, italic=False):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.italic = italic

title = doc.add_heading('', level=0)
r = title.add_run(
    'C\u00c2U CHUY\u1ec6N TRUY\u1ec0N C\u1ea2M H\u1ee8NG\n'
    'TS. Nguy\u1ec5n Th\u1ecb Ph\u01b0\u01a1ng Th\u1ea3o\n'
    'Ch\u1ee7 t\u1ecbch H\u0110QT Vietjet'
)
r.font.name = 'Times New Roman'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p('Tr\u00edch t\u1eeb cu\u1ed9c h\u1ecdp x\u00e2y d\u1ef1ng ch\u01b0\u01a1ng tr\u00ecnh \u0111\u00e0o t\u1ea1o l\u00e3nh \u0111\u1ea1o \u2014 H\u1ed9i nh\u1eadp & T\u1ecfa s\u00e1ng.', italic=True)
doc.add_paragraph()

# ============================
h('M\u1ed7i con ng\u01b0\u1eddi \u0111\u1ec1u c\u00f3 th\u1ec3 tr\u1edf n\u00ean v\u0129 \u0111\u1ea1i')
p('Tr\u01b0\u1edbc khi \u0111i v\u00e0o t\u1eebng c\u00e2u chuy\u1ec7n, Ch\u1ee7 t\u1ecbch \u0111\u1eb7t ra m\u1ed9t th\u00f4ng \u0111i\u1ec7p xuy\u00ean su\u1ed1t:')
p('"M\u1ed7i con ng\u01b0\u1eddi c\u1ee7a ch\u00fang ta \u0111\u1ec1u c\u00f3 th\u1ec3 tr\u1edf n\u00ean v\u0129 \u0111\u1ea1i. L\u00e0m nh\u1eefng c\u00e1i vi\u1ec7c l\u1edbn c\u00f3 \u00fd ngh\u0129a, m\u00e0 t\u1eeb nh\u1eefng c\u00e1i ho\u1ea1t \u0111\u1ed9ng r\u1ea5t l\u00e0 \u0111\u01a1n gi\u1ea3n, gi\u1ea3n d\u1ecb c\u1ee7a m\u00ecnh th\u00f4i."', italic=True)
doc.add_paragraph()

# 1
h('1. C\u00e1i h\u1ed3n c\u1ee7a \u0111\u00e0o t\u1ea1o')
p('Ch\u1ee7 t\u1ecbch nh\u1eafc l\u1ea1i r\u1eb1ng tinh th\u1ea7n c\u1ee7a ch\u01b0\u01a1ng tr\u00ecnh \u0111\u00e0o t\u1ea1o Vietjet v\u1ed1n c\u00f3 m\u1ed9t \u0111i\u1ec1u r\u1ea5t kh\u00e1c bi\u1ec7t. \u0110\u00f3 l\u00e0 c\u00e1i m\u00e0 Ch\u1ee7 t\u1ecbch g\u1ecdi l\u00e0 H\u1ed9i nh\u1eadp T\u1ecfa s\u00e1ng \u2014 kh\u00e1c v\u1edbi nh\u1eefng ho\u1ea1t \u0111\u1ed9ng \u0111\u00e0o t\u1ea1o th\u00f4ng th\u01b0\u1eddng. Nh\u01b0ng theo th\u1eddi gian, c\u00e1i tinh th\u1ea7n \u1ea5y \u0111\u00e3 d\u1ea7n mai m\u1ed9t.')
p('"N\u00f3 thi\u1ebfu c\u00e1i ki\u1ec3u gi\u1ed1ng nh\u01b0 con ng\u01b0\u1eddi, n\u00f3 thi\u1ebfu c\u00e1i h\u1ed3n c\u1ee7a n\u00f3. Ch\u00ednh nh\u1eefng c\u00e1i tinh th\u1ea7n \u0111\u1ea5y n\u00f3 m\u1edbi l\u00e0 t\u1ea1o cho ng\u01b0\u1eddi ta c\u00e1i \u0111\u1ed9ng l\u1ef1c."')
p('Ch\u1ee7 t\u1ecbch mu\u1ed1n quay tr\u1edf l\u1ea1i c\u00e1i kh\u1edf \u0111\u1ea7u \u0111\u00f3 \u2014 n\u01a1i m\u00e0 \u0111\u00e0o t\u1ea1o kh\u00f4ng ch\u1ec9 l\u00e0 k\u1ecbch b\u1ea3n, m\u00e0 l\u00e0 n\u01a1i ng\u01b0\u1eddi ta t\u00ecm th\u1ea5y gi\u00e1 tr\u1ecb c\u1ee7a b\u1ea3n th\u00e2n trong t\u1eadp th\u1ec3, n\u01a1i m\u00e0 khi g\u1eb7p kh\u00f3 kh\u0103n, m\u1ecdi ng\u01b0\u1eddi t\u1eadp trung v\u00e0o c\u00e1ch v\u01b0\u1ee3t qua v\u00e0 ph\u1ed1i h\u1ee3p \u0111\u1ed9i nh\u00f3m.')
doc.add_paragraph()

# 2
h('2. H\u00ecnh \u1ea3nh t\u1ed5 bay ng\u00e0y x\u01b0a v\u00e0 b\u00e2y gi\u1edd')
p('Ch\u1ee7 t\u1ecbch k\u1ec3 l\u1ea1i h\u00ecnh \u1ea3nh ng\u00e0y x\u01b0a: t\u1ed5 bay briefing t\u1ea1i tr\u1ee5 s\u1edf c\u00f4ng ty, x\u1ebfp h\u00e0ng ngay ng\u1eafn, ki\u1ec3m tra l\u1ea1i v\u1edbi nhau, r\u1ed3i di\u1ec5u trong v\u0103n ph\u00f2ng. Nh\u00e2n vi\u00ean \u0111\u1ea7u ti\u00ean nh\u00ecn t\u1ed5 bay chu\u1ea9n b\u1ecb l\u00e0m nhi\u1ec7m v\u1ee5, h\u1ecd t\u1ef1 h\u00e0o, h\u1ecd th\u1ea5y gi\u00e1 tr\u1ecb v\u00e0 \u00fd ngh\u0129a c\u1ee7a c\u00f4ng vi\u1ec7c.')
p('"Ch\u1ee9 c\u00f2n b\u00e2y gi\u1edd l\u00e0 c\u1ee9 \u0111\u1ea7u t\u1eaft m\u1eb7t t\u1ed1i, ra kh\u1ecfi nh\u00e0 2-3 gi\u1edd s\u00e1ng, ra s\u00e2n bay r\u1ed3i v\u1ec1 \u0111i ng\u1ee7. Ki\u1ebfm m\u1ed9t s\u1ed1 ti\u1ec1n th\u00ec kh\u00f4ng nh\u1ecf, nh\u01b0ng m\u00e0 c\u0169ng \u0111\u1ee7 s\u1ed1ng. Kh\u00f4ng c\u00f3 th\u1eddi gian ti\u00eau."')
p('Ng\u01b0\u1eddi ta r\u1eddi \u0111i kh\u00f4ng ph\u1ea3i v\u00ec ch\u1ebf \u0111\u1ed9 \u2014 ch\u1ebf \u0111\u1ed9 t\u1ed1t. Nh\u01b0ng c\u00e1i m\u00f4i tr\u01b0\u1eddng l\u00e0m vi\u1ec7c khi\u1ebfn ng\u01b0\u1eddi ta kh\u00f4ng c\u00f2n c\u1ea3m th\u1ea5y \u00fd ngh\u0129a.')
doc.add_paragraph()

# 3
h('3. Nh\u00e2n qu\u1ea3')
p('Ch\u1ee7 t\u1ecbch chia s\u1ebb tri\u1ebft l\u00fd qu\u1ea3n l\u00fd c\u1ee7a m\u00ecnh:')
p('"M\u00ecnh tin v\u00e0o nh\u00e2n qu\u1ea3. Nh\u00e2n n\u00e0o qu\u1ea3 \u0111\u1ea5y. M\u00ecnh l\u00e0m ch\u01b0a t\u1ed1t th\u00ec cu\u1ed9c s\u1ed1ng s\u1ebd c\u00f3 c\u00e1i qu\u1ea3 t\u01b0\u01a1ng \u1ee9ng. Ch\u1ee9 kh\u00f4ng c\u1ea7n ph\u1ea3i c\u00f3 h\u00e0nh \u0111\u1ed9ng g\u00ec x\u1eed l\u00fd c\u0103ng th\u1eb3ng."')
p('"Cu\u1ed9c s\u1ed1ng, ng\u01b0\u1eddi ta l\u00e0 con ng\u01b0\u1eddi, ng\u01b0\u1eddi ta c\u00f3 quy\u1ec1n sai. Th\u00ec m\u00ecnh h\u00e3y t\u1ea1o cho ng\u01b0\u1eddi ta c\u01a1 h\u1ed9i \u0111\u1ec3 c\u00f3 th\u1ec3 kh\u1eafc ph\u1ee5c."')
doc.add_paragraph()

# 4
h('4. C\u00f4ng b\u1eb1ng')
p('"T\u1ea1o n\u00ean s\u1ef1 c\u00f4ng b\u1eb1ng trong t\u1ed5 ch\u1ee9c l\u00e0 c\u1ea7n th\u1eddi gian, c\u1ea7n c\u01a1 ch\u1ebf. Ch\u1ee9 kh\u00f4ng ph\u1ea3i b\u00e2y gi\u1edd c\u1ea5p d\u01b0\u1edbi \u0111\u1ee5ng ch\u1ea1m m\u1ed9t c\u00e1i g\u00ec \u0111\u1ea5y m\u00e0 m\u00ecnh cho r\u1eb1ng l\u00e0 sai, th\u00ec m\u00ecnh x\u1eed l\u00fd r\u1ea5t l\u00e0 m\u1ea1nh. Trong khi \u0111\u00f3 l\u00e3nh \u0111\u1ea1o \u1edf tr\u00ean kh\u00f4ng g\u01b0\u01a1ng m\u1eabu, th\u00ec \u0111\u1ea5y n\u00f3 kh\u00f4ng c\u00f4ng b\u1eb1ng."')
doc.add_paragraph()

# 5
h('5. N\u00e2ng c\u00f4ng su\u1ea5t s\u00e2n bay T\u00e2n S\u01a1n Nh\u1ea5t')
p('N\u0103m 2013, Vietjet b\u1eaft tay v\u00e0o d\u1ef1 \u00e1n n\u00e2ng cao n\u0103ng l\u1ef1c \u0111i\u1ec1u h\u00e0nh c\u1ea5t h\u1ea1 c\u00e1nh \u1edf s\u00e2n bay T\u00e2n S\u01a1n Nh\u1ea5t. N\u0103m 2014 b\u1eaft \u0111\u1ea7u ch\u1ea1y. C\u00f4ng su\u1ea5t \u0111\u01b0\u1ee3c n\u00e2ng l\u00ean g\u1ea5p \u0111\u00f4i \u2014 t\u1eeb 30 l\u01b0\u1ee3t c\u1ea5t h\u1ea1 c\u00e1nh l\u00ean 60 l\u01b0\u1ee3t. Hi\u1ec7n t\u1ea1i \u0111ang v\u1eadn h\u00e0nh 48 l\u01b0\u1ee3t. L\u01b0\u1ee3ng kh\u00e1ch qua T\u00e2n S\u01a1n Nh\u1ea5t t\u0103ng t\u1eeb 24 tri\u1ec7u l\u00ean 42 tri\u1ec7u l\u01b0\u1ee3t m\u1ed7i n\u0103m.')
p('"D\u1ef1 \u00e1n \u1ea5y l\u00e0m cho c\u1ea3 ng\u00e0nh h\u00e0ng kh\u00f4ng. ACV m\u01b0\u1eddi m\u1ea5y n\u0103m li\u1ec1n, m\u1ed7i n\u0103m \u0111\u00f3n 16 tri\u1ec7u l\u01b0\u1ee3t kh\u00e1ch, l\u1ee3i nhu\u1eadn m\u1ea5y tr\u0103m tri\u1ec7u \u0111\u00f4 la m\u1ed9t n\u0103m. Ch\u1ee9 kh\u00f4ng ph\u1ea3i l\u00e0 doanh s\u1ed1 \u0111\u00e2u. Th\u00ec l\u00e0m sao m\u00e0 ch\u1eb3ng c\u00f3 ti\u1ec1n \u0111i x\u00e2y s\u00e2n bay."')
doc.add_paragraph()

# 6
h('6. Gi\u1ea3i c\u1ee9u s\u00e0n ch\u1ee9ng kho\u00e1n')
p('Khi s\u00e0n ch\u1ee9ng kho\u00e1n Vi\u1ec7t Nam b\u1ecb treo, Vietjet \u0111\u00e3 tham gia gi\u1ea3i c\u1ee9u cho c\u1ea3 th\u1ecb tr\u01b0\u1eddng v\u1ed1n. Ch\u1ee7 t\u1ecbch so s\u00e1nh: c\u1ea3 n\u01b0\u1edbc M\u1ef9 c\u00f3 s\u00e0n New York v\u00e0 Nasdaq \u2014 2 c\u00e1i s\u00e0n ch\u1ee9ng kho\u00e1n. C\u1ea3 n\u01b0\u1edbc Anh c\u00f3 London Stock Exchange.')
p('"C\u00f3 h\u01a1n 100 c\u00f4ng ty ch\u1ee9ng kho\u00e1n, 4 tri\u1ec7u kh\u00e1ch h\u00e0ng l\u00e0 t\u1ed5 ch\u1ee9c v\u00e0 c\u00e1 nh\u00e2n trong n\u01b0\u1edbc v\u00e0 ngo\u00e0i n\u01b0\u1edbc. M\u1ed7i ng\u00e0y giao d\u1ecbch kho\u1ea3ng 1 t\u1ef7 \u0111\u00f4 la, m\u00e0 cao l\u00e0 2 t\u1ef7 \u0111\u00f4 la. M\u00ecnh c\u0169ng nh\u1ea3y v\u00e0o gi\u1ea3i c\u1ee9u."')
p('"Bao nhi\u00eau ng\u01b0\u1eddi \u0111\u01b0\u1ee3c h\u01b0\u1edfng l\u1ee3i t\u1eeb \u0111\u00f3. Ch\u00ednh ph\u1ee7 th\u00ec \u0111\u01b0\u1ee3c c\u00e1i s\u00e0n, \u0111\u01b0\u1ee3c c\u1ea3 n\u01b0\u1edbc. C\u00e1c c\u00f4ng ty, c\u00e1c nh\u00e0 \u0111\u1ea7u t\u01b0 th\u00ec \u0111\u01b0\u1ee3c h\u01b0\u1edfng l\u1ee3i. B\u1ea3n th\u00e2n c\u00e1c c\u00f4ng ty ni\u00eam y\u1ebft \u1edf \u0111\u1ea5y th\u00ec ng\u01b0\u1eddi ta c\u00f3 th\u1ecb tr\u01b0\u1eddng v\u1ed1n \u0111\u1ec3 huy \u0111\u1ed9ng v\u1ed1n, ho\u1ea1t \u0111\u1ed9ng, ph\u00e1t tri\u1ec3n. Th\u1ebf c\u00f2n m\u00ecnh th\u00ec kh\u00f4ng c\u1ea7n c\u00e1i g\u00ec c\u1ea3."')
doc.add_paragraph()

# 7
h('7. Ngo\u1ea1i giao qua mua m\u00e1y bay')
p('M\u1ed7i h\u1ee3p \u0111\u1ed3ng mua m\u00e1y bay l\u00e0 m\u1ed9t c\u01a1 h\u1ed9i x\u00e2y d\u1ef1ng quan h\u1ec7 ngo\u1ea1i giao kinh t\u1ebf.')
p('"Mua m\u00e1y bay Boeing l\u00e0 \u0111i v\u1edbi n\u01b0\u1edbc M\u1ef9 \u2014 \u00f4ng Trump, \u00f4ng Obama, \u00f4ng Bush. Mua m\u00e1y bay Airbus l\u00e0 \u0111i v\u1edbi Ph\u00e1p, \u0110\u1ee9c, T\u00e2y Ban Nha \u2014 nh\u1eefng n\u01b0\u1edbc s\u1ea3n xu\u1ea5t Airbus. Mua m\u00e1y bay COMAC c\u1ee7a Trung Qu\u1ed1c l\u00e0 ch\u01b0\u01a1ng tr\u00ecnh c\u1ee7a \u00f4ng T\u1eadp C\u1eadn B\u00ecnh."')
p('"L\u00e0 m\u00ecnh t\u1ea1o n\u00ean x\u00e2y d\u1ef1ng m\u1ed1i quan h\u1ec7 ngo\u1ea1i giao, x\u00e2y d\u1ef1ng kinh t\u1ebf gi\u1eefa c\u00e1c qu\u1ed1c gia, d\u00e2n t\u1ed9c. M\u00ecnh s\u1ebd gi\u1ea3m b\u1edbt \u0111i nh\u1eefng chi\u1ebfn tranh, nh\u1eefng s\u1ef1 xung \u0111\u1ed9t."')
doc.add_paragraph()

# 8
h('8. APEC v\u00e0 n\u0103ng l\u1ef1c tri\u1ec3n khai th\u1ea7n t\u1ed1c')
p('Vietjet ph\u1ee5c v\u1ee5 APEC 2 l\u1ea7n, n\u0103m 2017 \u2014 "\u0110\u01b0a Vi\u1ec7t Nam v\u00e0o th\u1ebf gi\u1edbi v\u00e0 \u0111\u01b0a th\u1ebf gi\u1edbi v\u1ec1 Vi\u1ec7t Nam."')
p('"C\u00e1i gi\u1ea3i ph\u00e1p chi\u1ebfu s\u00e1ng m\u00ecnh l\u00e0m c\u00f3 90 ng\u00e0y th\u00f4i \u2014 trong khi \u0111\u00f3 H\u00e0n Qu\u1ed1c tri\u1ec3n khai 12 n\u0103m m\u00e0 kh\u00f4ng ch\u1ea1y, \u0111\u1ebfn l\u00fac b\u00e0n giao \u0111\u00e3 17 n\u0103m. X\u00e2y cung h\u1ed9i ngh\u1ecb 4 th\u00e1ng."')
p('"Ng\u00e0y 9 th\u00e1ng 1 n\u0103m ngo\u00e1i, kh\u00f4ng ai m\u1eddi m\u00ecnh \u0111\u00e2u, nh\u01b0ng m\u00e0 th\u1ea5y th\u1eddi cu\u1ed9c n\u00f3 c\u0103ng th\u1eb3ng. M\u00ecnh c\u00f3 m\u1ed1i quan h\u1ec7 t\u1ed1t v\u1edbi \u00f4ng Trump, \u0110\u1ea1i s\u1ee9 qu\u00e1n M\u1ef9 m\u1eddi m\u00ecnh, th\u00ec m\u00ecnh bay lu\u00f4n, mang m\u00e1y bay \u0111i lu\u00f4n. C\u0169ng t\u1ea1o n\u00ean thi\u1ec7n c\u1ea3m, c\u0169ng \u0111\u1ee1 c\u0103ng th\u1eb3ng."')
doc.add_paragraph()

# 9
h('9. Thi\u1ec7n nguy\u1ec7n COVID v\u00e0 tri\u1ebft l\u00fd thi\u1ebft k\u1ebf')
p('"M\u00f9a Covid m\u00ecnh lao \u0111i ch\u1edf v\u1eafc-xin, thi\u1ebft b\u1ecb y t\u1ebf, bao nhi\u00eau nh\u1eefng bao g\u1ea1o. Kh\u00f4ng c\u00f3 nh\u1eefng c\u00e1i \u0111\u00f3 th\u00ec l\u00e0m sao m\u00e0 ng\u01b0\u1eddi ta \u0111\u01b0\u1ee3c y\u00ean l\u00f2ng."')
p('Ch\u1ee7 t\u1ecbch k\u1ec3 ti\u1ebfp v\u1ec1 tri\u1ebft l\u00fd thi\u1ebft k\u1ebf:')
p('"T\u1ea5t c\u1ea3 nh\u1eefng c\u00e1i g\u00ec m\u00e0 m\u00ecnh l\u00e0m \u2014 thi\u1ebft k\u1ebf m\u1ed9t c\u00e1i nh\u00e0 chung c\u01b0, l\u00e0 ph\u1ea3i c\u00f3 c\u00e1i thang m\u00e1y l\u1edbn \u0111\u1ec3 d\u00e0nh cho c\u00e1ng c\u1ee9u th\u01b0\u01a1ng v\u00e0 quan t\u00e0i. T\u1eeb c\u00e1i thi\u1ebft k\u1ebf l\u00e0 m\u00ecnh c\u0169ng ngh\u0129 \u0111\u1ebfn con ng\u01b0\u1eddi \u2014 l\u00fac ng\u01b0\u1eddi ta b\u1ec7nh, ng\u01b0\u1eddi ta c\u1ea7n ph\u1ea3i \u0111\u01b0a \u0111i c\u00e1ng. \u1ede \u0111\u00e2u c\u00f3 b\u1eadc thang l\u00e0 m\u00ecnh ph\u1ea3i c\u00f3 c\u00e1i \u0111\u01b0\u1eddng d\u1ed1c cho xe n\u00f4i, cho xe ng\u01b0\u1eddi khuy\u1ebft t\u1eadt. \u0110\u1ea5y l\u00e0 tinh th\u1ea7n m\u00ecnh \u0111\u01b0a v\u00e0o m\u1ecdi ho\u1ea1t \u0111\u1ed9ng."')
doc.add_paragraph()

# 10
h('10. \u0110\u00f3ng g\u00f3p to\u00e0n di\u1ec7n')
p('"C\u1edd vua l\u00ean top 7 th\u1ebf gi\u1edbi. B\u00f3ng \u0111\u00e1 nam v\u00e0o World Cup, b\u00f3ng \u0111\u00e1 n\u1eef c\u0169ng v\u00e0o World Cup. To\u00e0n di\u1ec7n cho \u0111\u1ea5t n\u01b0\u1edbc, \u1edf t\u1ea5t c\u1ea3 c\u00e1c l\u0129nh v\u1ef1c \u2014 v\u0103n h\u00f3a, x\u00e3 h\u1ed9i, kinh t\u1ebf, t\u00e0i ch\u00ednh."')
doc.add_paragraph()

# 11
h('11. Tinh th\u1ea7n coaching')
p('"M\u1ed9t trong nh\u1eefng tinh th\u1ea7n m\u00ecnh mang \u0111\u1ebfn \u0111\u00e2y v\u1edbi c\u00e1c l\u00e3nh \u0111\u1ea1o l\u00e0 tinh th\u1ea7n h\u1ecdc th\u1eadt. C\u00e1i tinh th\u1ea7n ph\u1ea3i tr\u00e2n tr\u1ecdng \u0111\u00e0o t\u1ea1o. M\u00e0 l\u00fac n\u00e0o \u0111\u00f3 n\u00f3 mai m\u1ed9t \u0111i th\u00ec quay tr\u1edf l\u1ea1i. V\u00e0 m\u1ed7i ng\u01b0\u1eddi ph\u1ea3i tr\u1edf th\u00e0nh ng\u01b0\u1eddi coach cho nh\u1eefng ng\u01b0\u1eddi xung quanh. C\u0169ng ph\u1ea3i bi\u1ebft c\u00e1ch h\u1ecdc h\u1ecfi."')
doc.add_paragraph()

# Kết
h('Th\u00f4ng \u0111i\u1ec7p')
p('"M\u1ed7i con ng\u01b0\u1eddi \u0111\u1ec1u c\u00f3 th\u1ec3 tr\u1edf n\u00ean v\u0129 \u0111\u1ea1i. L\u00e0m nh\u1eefng vi\u1ec7c l\u1edbn c\u00f3 \u00fd ngh\u0129a, m\u00e0 t\u1eeb nh\u1eefng ho\u1ea1t \u0111\u1ed9ng r\u1ea5t l\u00e0 \u0111\u01a1n gi\u1ea3n, gi\u1ea3n d\u1ecb c\u1ee7a m\u00ecnh th\u00f4i."', italic=True)

doc.save('C\u00e2u chuy\u1ec7n Madame Th\u1ea3o.docx')
print('Done')
