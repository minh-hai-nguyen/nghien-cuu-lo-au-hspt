# -*- coding: utf-8 -*-
"""Apply Heading styles v2 - context-aware (MỞ ĐẦU vs CHƯƠNG vs KẾT LUẬN).
KHOI PHUC tu backup truoc khi apply.
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
BACKUP_BEFORE_HEADING = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026_BEFORE_APPLY_HEADING.docx')

# Restore first
print(f"Restoring from: {BACKUP_BEFORE_HEADING}")
shutil.copy(BACKUP_BEFORE_HEADING, FILE)

d = Document(FILE)


def in_table(p):
    """True if paragraph is inside a table."""
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag in (qn('w:tbl'), qn('w:tc')):
            return True
        parent = parent.getparent()
    return False


# H1 keywords (whole text equals)
H1_EXACT = {
    'MỞ ĐẦU', 'KẾT LUẬN VÀ KIẾN NGHỊ',
    'TÀI LIỆU THAM KHẢO', 'PHỤ LỤC', 'MỤC LỤC',
    'LỜI CAM ĐOAN', 'LỜI CẢM ƠN',
    'DANH MỤC CÔNG TRÌNH ĐÃ CÔNG BỐ',
    'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
    'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
}
# H1 strict prefix: must be "CHƯƠNG <digit/Roman>"
CHUONG_RE = re.compile(r'^CHƯƠNG\s+[\dIVX]+\s*[:.]?\s*$', re.IGNORECASE)
TIEU_KET_RE = re.compile(r'^TIỂU KẾT CHƯƠNG\s+[\dIVX]+\s*$', re.IGNORECASE)

# Chapter title (caps, structural, follows CHƯƠNG): allow it as H1
CHAPTER_TITLE_KEYWORDS = ['CƠ SỞ LÝ LUẬN', 'TỔ CHỨC VÀ PHƯƠNG PHÁP',
                          'KẾT QUẢ NGHIÊN CỨU', 'PHÂN TÍCH KẾT QUẢ',
                          'ĐỀ XUẤT', 'CHƯƠNG TRÌNH']


def determine_section_context(prev_h1_text):
    """Given the most recent H1 text, return one of:
    - 'mo_dau' (MỞ ĐẦU, KẾT LUẬN VÀ KIẾN NGHỊ)
    - 'chuong' (CHƯƠNG X + title)
    - 'other' (cover, danh muc, etc.)
    """
    if not prev_h1_text:
        return 'other'
    upper = prev_h1_text.upper().strip()
    if upper.startswith('MỞ ĐẦU') or 'KẾT LUẬN' in upper:
        return 'mo_dau'
    if upper.startswith('CHƯƠNG ') or any(kw in upper for kw in CHAPTER_TITLE_KEYWORDS):
        return 'chuong'
    if upper.startswith('TIỂU KẾT'):
        return 'chuong'  # still in chương
    return 'other'


def count_numbers_in_prefix(text):
    """Count how many numbers are in the leading numbering pattern.
    E.g., '1.1.5.1.' -> 4, '7.1 ' -> 2, '1. ' -> 1, '1.1 ' -> 2.
    Returns None if no numbering prefix."""
    m = re.match(r'^(\d+(?:\.\d+)*)\.?\s', text)
    if not m:
        return None
    nums = m.group(1).split('.')
    return len(nums)


def decide_heading(text, p, context):
    """Return level or None."""
    text = text.strip()
    if not text or len(text) > 250:
        return None
    # Skip captions
    if re.match(r'^(Bảng|Hình|Biểu đồ|Sơ đồ)\s+\d', text, re.IGNORECASE):
        return None
    # Skip sub-items 'a/', 'b/', 'c/'
    if re.match(r'^[a-z]/\s', text):
        return None

    upper = text.upper()

    # H1 exact
    if upper in H1_EXACT:
        return 1

    # CHƯƠNG <num> only (not 'Chương trình', 'Chương này')
    if CHUONG_RE.match(text):
        return 1

    # TIỂU KẾT CHƯƠNG <num>
    if TIEU_KET_RE.match(text):
        return 1

    # Chapter title (all caps, contains key word)
    if upper == text and len(text) < 200:
        if any(kw in upper for kw in CHAPTER_TITLE_KEYWORDS):
            return 1

    # Numbered patterns
    n_nums = count_numbers_in_prefix(text)
    if n_nums is None:
        # Special: "Giả thuyết X" patterns
        if re.match(r'^Giả thuyết\s+[1-9]\s*[—–-]', text):
            return 3
        return None

    if context == 'mo_dau':
        # 1 num = H2; 2 nums = H3; 3 nums = H4
        return min(n_nums + 1, 5)
    elif context == 'chuong':
        # 2 nums = H2; 3 nums = H3; 4 nums = H4
        if n_nums < 2:
            return None  # unusual, skip
        return min(n_nums, 5)
    else:
        # Other contexts (cover, danh muc) - skip numbered items
        return None


# ============================================================
# SCAN with context tracking
# ============================================================
print("\n=== SCANNING WITH CONTEXT ===")
prev_h1_text = ''  # tracks current section context
context = 'other'

proposed = []
for i, p in enumerate(d.paragraphs):
    if in_table(p):
        continue
    text = p.text.strip()
    if not text:
        continue

    # Skip cover/first 60 paras
    if i < 60:
        continue

    current_style = p.style.name if p.style else 'Normal'

    # First, see if THIS paragraph is itself H1 — if so, update context
    upper = text.upper()
    is_h1 = (upper in H1_EXACT or CHUONG_RE.match(text) or TIEU_KET_RE.match(text)
             or (upper == text and len(text) < 200 and any(kw in upper for kw in CHAPTER_TITLE_KEYWORDS)))
    if is_h1:
        prev_h1_text = text
        context = determine_section_context(prev_h1_text)

    level = decide_heading(text, p, context)
    if level is None:
        continue
    expected_styles = [f'Heading {level}', f'Tiêu đề {level}']
    if current_style in expected_styles:
        continue
    proposed.append({
        'idx': i,
        'text': text[:100],
        'current_style': current_style,
        'target_level': level,
        'context': context,
    })

print(f"\nProposed {len(proposed)} paragraphs to apply Heading style")
print()
print(f"{'idx':>5} {'L':>2} {'ctx':>8} {'cur_style':>15}  text")
print("-"*120)
for prop in proposed:
    print(f"{prop['idx']:>5} {prop['target_level']:>2} {prop['context']:>8} {prop['current_style'][:15]:>15}  {prop['text']}")

# ============================================================
# APPLY
# ============================================================
print()
print("=== APPLYING ===")
applied = 0
failed = 0
for prop in proposed:
    p = d.paragraphs[prop['idx']]
    target = f"Heading {prop['target_level']}"
    try:
        p.style = d.styles[target]
        applied += 1
    except Exception as e:
        try:
            p.style = d.styles[f"Tiêu đề {prop['target_level']}"]
            applied += 1
        except Exception:
            failed += 1
            print(f"  FAIL para {prop['idx']}: {e}")

print(f"Applied: {applied}, Failed: {failed}")
d.save(FILE)
print(f"Saved: {FILE}")
