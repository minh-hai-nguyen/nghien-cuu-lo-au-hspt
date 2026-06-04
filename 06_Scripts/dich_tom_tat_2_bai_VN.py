# -*- coding: utf-8 -*-
"""
Tao 2 ban dich + 2 tom tat cho 2 bai VN moi:
- Bai 59: TCNCYH 2025 (Dao Thi Ngoan, ĐH Y Ha Noi, SV nam 4, n=196 DASS-21)
- Bai 60: Duong et al. 2025 (Soc Psychiatry Q1, n=2.631 HS THPT TPHCM)
"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
DICH_DIR = os.path.join(ROOT, '03_Ban-dich')
PAGE_W = 16.0

# ===== Helpers =====
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w * 567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
def P(doc, text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
def save(doc, path):
    doc.save(path)
    d = Document(path); chars = sum(len(x.text) for x in d.paragraphs)
    print(f'  {os.path.basename(path)}: {chars} chars, {len(d.tables)} tables')

# ============================================================
# BÀI 59 — Đào Thị Ngoãn et al. 2025 — Bản dịch (full document)
# ============================================================
print('\n[59] Bài 59 — TCNCYH 2025 — bản dịch + tóm tắt')

# 59.1 — Bản dịch (file thực tế là tiếng Việt nên đây là bản chuẩn hoá format)
doc = make_doc()
P(doc, 'Link bài báo gốc: TCNCYH 187 (02) - 2025, Tạp chí Nghiên cứu Y học, tr. 296-303', size=10)
H(doc, 'Thực trạng tâm lý của sinh viên năm thứ tư Trường Đại học Y Hà Nội năm học 2023-2024', level=1)

H(doc, 'THÔNG TIN THƯ MỤC', 2)
table(doc,
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Đào Thị Ngoãn, Phạm Tùng Sơn, Trần Thị Ngọc Trang, Nguyễn Thị Nguyệt, Võ Trung Hiếu, Phó Tuấn Vinh, Trần Hồng Thái, Hoàng Bảo An, Trần Phạm Mạnh, Phan Thị Minh Ngọc'],
        ['Đơn vị', 'Trường Đại học Y Hà Nội'],
        ['Tạp chí', 'Tạp chí Nghiên cứu Y học (TCNCYH), tập 187, số 02-2025'],
        ['Trang', '296–303'],
        ['Tác giả liên hệ', 'Đào Thị Ngoãn — daothingoan@hmu.edu.vn'],
        ['Ngày nhận', '19/11/2024'],
        ['Ngày chấp nhận', '24/12/2024'],
        ['Loại NC', 'Mô tả cắt ngang'],
        ['Mẫu', '196 sinh viên năm thứ 4 Trường ĐH Y Hà Nội (126 ngành bác sĩ + 70 cử nhân)'],
        ['Công cụ', 'DASS-21 + Bộ câu hỏi nhu cầu tham vấn tâm lý (Viện ĐTYHDP&YTCC, ĐH Y Hà Nội)'],
        ['Thời gian', '01/2024 – 12/2024'],
    ],
    widths=[3.5, 12.0])

H(doc, 'TÓM TẮT', 2)
P(doc, 'Sinh viên ngành y có nhiều áp lực do đặc thù về thời gian đào tạo kéo dài, khối lượng kiến '
       'thức lớn. Tại Trường Đại học Y Hà Nội, năm thứ 4 là thời điểm sinh viên bắt đầu học thực '
       'hành lâm sàng tại bệnh viện, phải đối mặt với nhiều mối quan hệ xã hội, gặp nhiều khó '
       'khăn hơn trong việc sắp xếp thời gian học tập và nghỉ ngơi, từ đó có thể phát sinh nhiều '
       'hơn các vấn đề về tâm lý. Nghiên cứu thực hiện với 2 mục tiêu: (1) Mô tả kết quả trắc '
       'nghiệm tâm lý DASS-21; (2) Đánh giá mối liên quan giữa kết quả DASS-21 với một số yếu tố. '
       '196 sinh viên năm thứ 4 tất cả các ngành được lựa chọn ngẫu nhiên, được hướng dẫn trả '
       'lời bộ câu hỏi về một số vấn đề trong cuộc sống, học tập; thực hiện thang đánh giá tâm lý '
       'DASS-21. Kết quả DASS-21 cho thấy có 33,67 % sinh viên bị trầm cảm, 43,88 % bị lo âu và '
       '31,63 % bị căng thẳng. 57,65 % sinh viên được khảo sát có mắc ít nhất 1 trong 3 rối loạn '
       'sức khoẻ tâm thần. 11,73 % sinh viên mắc ít nhất 1 rối loạn từ mức nghiêm trọng trở lên. '
       'Có một số mối liên quan giữa kết quả DASS-21 và các vấn đề gặp phải trong đời sống và '
       'học tập như điểm tích luỹ, tài chính, giấc ngủ, ngành học, sự hài lòng về học tập và các '
       'mối quan hệ xã hội.')

H(doc, 'PHƯƠNG PHÁP', 2)
P(doc, 'Thiết kế: Mô tả cắt ngang. Cỡ mẫu tính theo công thức cho một tỷ lệ quần thể với p = 0,85 '
       '(tỷ lệ tham khảo từ NC trước về sinh viên Y mắc rối loạn tâm thần), d = 0,05, α = 0,05 → '
       'n_min = 195. Thực tế khảo sát: 196 sinh viên (cao hơn yêu cầu).')
P(doc, 'Cách chọn mẫu: Chọn mẫu ngẫu nhiên đơn theo danh sách sinh viên năm 4 các ngành học '
       '(4 ngành hệ bác sĩ + 6 ngành hệ cử nhân) — danh sách quản lý bởi Phòng CTHVSV&QLKTX. '
       'Dùng lệnh RAND của Excel sinh giá trị ngẫu nhiên cho từng sinh viên, sắp xếp tăng dần và '
       'lấy từ trên xuống. Số mẫu mỗi ngành xác định theo tương quan giữa số sinh viên ngành đó '
       'trên tổng 901 sinh viên toàn khối với cỡ mẫu mục tiêu 195.')
P(doc, 'Công cụ: (1) Bộ câu hỏi nhu cầu tham vấn tâm lý do Viện ĐTYHDP&YTCC, ĐH Y Hà Nội phát '
       'triển năm 2019, gồm thông tin chung (ngành học, giới, BMI, nơi ở, đối tượng sống cùng, '
       'kết quả học tập tích luỹ) và thực trạng khó khăn ở các khía cạnh học tập, sinh hoạt, tài '
       'chính, mối quan hệ xã hội. (2) Bộ DASS-21 đã chuẩn hoá bằng tiếng Việt bởi Viện Sức '
       'khoẻ Tâm thần, Bệnh viện Bạch Mai.')
P(doc, 'Quy trình: Sinh viên ký cam kết tự nguyện, trả lời tự điền tại phòng LAB Bộ môn Sinh lý '
       'học (tầng 1 nhà B1, ĐH Y Hà Nội). Nhập liệu Excel, phân tích STATA 15.0. Thống kê mô tả '
       '(tỷ lệ %), phân tích χ²/Fisher để so sánh khác biệt giữa các tỷ lệ, ngưỡng p < 0,05. Với '
       'liên quan có ý nghĩa thống kê, tính tỷ suất chênh OR.')

H(doc, 'KẾT QUẢ', 2)

H(doc, 'Bảng 1. Phân loại DASS-21 theo mức độ (n = 196)', 3)
table(doc,
    ['Mức độ', 'Trầm cảm n (%)', 'Lo âu n (%)', 'Căng thẳng n (%)'],
    [
        ['Bình thường', '130 (66,33 %)', '110 (56,12 %)', '134 (68,37 %)'],
        ['Nhẹ', '35 (17,86 %)', '31 (15,82 %)', '28 (14,29 %)'],
        ['Trung bình', '22 (11,22 %)', '39 (19,90 %)', '25 (12,76 %)'],
        ['Nghiêm trọng', '7 (3,57 %)', '9 (4,59 %)', '6 (3,06 %)'],
        ['Rất nghiêm trọng', '2 (1,02 %)', '7 (3,57 %)', '3 (1,53 %)'],
        ['Có rối loạn (≥ nhẹ)', '66 (33,67 %)', '86 (43,88 %)', '62 (31,63 %)'],
        ['Từ trung bình trở lên', '15,81 %', '28,06 %', '17,35 %'],
    ],
    widths=[5.0, 3.5, 3.5, 3.5])

H(doc, 'Bảng 2. Phân loại theo số lượng và mức độ rối loạn (n = 196)', 3)
table(doc,
    ['Phân loại', 'Số lượng', 'Tỷ lệ (%)'],
    [
        ['Không mắc rối loạn nào', '83', '42,35'],
        ['Mắc ít nhất 1 trong 3 rối loạn', '113', '57,65'],
        ['Mắc ít nhất 1 rối loạn từ mức nghiêm trọng trở lên', '23', '11,73'],
    ],
    widths=[8.5, 3.5, 3.5])

H(doc, 'Bảng 3. Yếu tố liên quan có ý nghĩa thống kê', 3)
table(doc,
    ['Outcome', 'Yếu tố nguy cơ', 'OR', 'p'],
    [
        ['Trầm cảm', 'Điểm tích luỹ Giỏi (so với Khá–TB)', '4,97', '0,005'],
        ['Lo âu', 'Hệ Bác sĩ (so với Cử nhân)', '1,86', '0,044'],
        ['Lo âu', 'Vấn đề giấc ngủ', '1,84', '0,038'],
        ['Lo âu', 'Điểm tích luỹ Giỏi (so với Khá–TB)', '4,69', '0,013'],
        ['Căng thẳng', 'Chưa hài lòng học tập (so với hài lòng)', '10,70', '0,005'],
        ['Căng thẳng', 'Vấn đề giấc ngủ', '2,90', '0,001'],
        ['Căng thẳng', 'Mối quan hệ xã hội có vấn đề', '2,92', '0,002'],
        ['Căng thẳng', 'Tài chính không đủ', '2,74', '0,012'],
        ['Căng thẳng', 'Điểm tích luỹ Giỏi', '3,82', '0,016'],
        ['≥1 rối loạn', 'Vấn đề giấc ngủ', '1,84', '0,037'],
        ['≥1 rối loạn', 'Mối quan hệ xã hội có vấn đề', '2,37', '0,024'],
    ],
    widths=[3.0, 7.0, 2.0, 2.0])

P(doc, 'Phát hiện đáng chú ý: Sinh viên đạt điểm tích luỹ MỨC GIỎI có tỷ lệ trầm cảm gấp 4,97 lần '
       'so với nhóm Khá – Trung bình (p = 0,005), tỷ lệ lo âu gấp 4,69 lần (p = 0,013), tỷ lệ căng '
       'thẳng gấp 3,82 lần (p = 0,016). Đây là HIỆN TƯỢNG NGHỊCH ĐẢO so với giả thuyết thông thường '
       '— sinh viên giỏi lại có nguy cơ cao nhất, gợi ý "perfectionism trap" / áp lực thành tích.')

H(doc, 'ĐẶC ĐIỂM ĐỐI TƯỢNG NGHIÊN CỨU', 2)
P(doc, '• 196 sinh viên: 126 ngành bác sĩ + 70 ngành cử nhân; nam 35,71 %, nữ 64,29 %.')
P(doc, '• BMI: bình thường/khoẻ mạnh 60,71 %; thiếu cân 24,49 %; béo phì 4,59 %.')
P(doc, '• Kết quả học tập: Giỏi 6,63 %, Khá 38,27 %, Trung bình khá 50,59 %, Trung bình 0,51 %.')
P(doc, '• Sự hài lòng học tập: Hài lòng 10,89 %; chưa hài lòng phương pháp 56,05 %; chưa hài lòng '
       'môi trường 11,69 %; kết quả không tương xứng cố gắng 21,37 %.')
P(doc, '• Tài chính: Đủ cho nhu cầu cơ bản 52 %; đủ cả giải trí 33 %; gặp khó khăn 15 %.')
P(doc, '• Giấc ngủ: 55,61 % gặp vấn đề về thời lượng/khả năng vào giấc ngủ/chất lượng giấc ngủ.')
P(doc, '• Mối quan hệ xã hội: 30,92 % gặp vấn đề về bạn bè, gia đình, tình yêu.')

H(doc, 'KẾT LUẬN', 2)
P(doc, 'Tỷ lệ rối loạn sức khoẻ tâm thần ở sinh viên năm 4 ĐH Y Hà Nội rất cao: 33,67 % trầm cảm, '
       '43,88 % lo âu, 31,63 % căng thẳng; 57,65 % mắc ít nhất 1 trong 3 rối loạn; 11,73 % mắc rối '
       'loạn từ mức nghiêm trọng trở lên. Các yếu tố liên quan có ý nghĩa thống kê: ngành học, '
       'điểm tích luỹ, sự hài lòng học tập, giấc ngủ, tài chính, mối quan hệ xã hội. Cần có chương '
       'trình hỗ trợ tâm lý cho sinh viên y, đặc biệt nhóm điểm cao có nguy cơ cao bất thường.')

# Phản biện CTH v5
H(doc, 'QUAN ĐIỂM PHẢN BIỆN', 2)
P(doc, 'Điểm mạnh:', bold=True, color=RGBColor(0xCC, 0, 0))
P(doc, '• Cỡ mẫu chọn ngẫu nhiên đơn từ danh sách 901 sinh viên — không phải mẫu thuận tiện. '
       'Đây là điểm khác biệt tích cực so với đa số NC SKTT sinh viên VN khác.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Sử dụng DASS-21 đã chuẩn hoá tiếng Việt (Viện SKTT Bạch Mai) — công cụ tin cậy.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Đo CẢ 3 trục (trầm cảm + lo âu + căng thẳng) — toàn diện hơn các NC chỉ đo 1 trục.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Phát hiện CỰC THÚ VỊ về điểm tích luỹ Giỏi → nguy cơ rất cao (OR = 4–5), gợi ý hiện '
       'tượng "perfectionism trap" trong môi trường y khoa cạnh tranh.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Đối tượng SV năm 4 vừa bắt đầu lâm sàng — thời điểm chuyển giao quan trọng, có giá '
       'trị triển khai can thiệp sớm.', color=RGBColor(0xCC, 0, 0))

P(doc, 'Hạn chế:', bold=True, color=RGBColor(0xCC, 0, 0))
P(doc, '• Đối tượng SINH VIÊN ĐẠI HỌC, không phải VTN học sinh THCS/THPT — không thể ngoại '
       'suy trực tiếp cho đề tài lo âu VTN VN.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Cỡ mẫu n = 196 — nhỏ cho phân tích đa biến tinh tế. Cỡ mẫu chỉ vừa đủ cho mô tả '
       '(theo công thức n_min = 195).', color=RGBColor(0xCC, 0, 0))
P(doc, '• Cắt ngang — không thể kết luận quan hệ nhân quả. Sinh viên giỏi có thể bị áp lực '
       'cao hoặc người bị áp lực cao mới đạt điểm giỏi (chiều ngược).', color=RGBColor(0xCC, 0, 0))
P(doc, '• Chỉ 1 trường (ĐH Y Hà Nội), 1 năm học cụ thể (Y4) — không đại diện toàn bộ SV y VN.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Tự báo cáo giấc ngủ, tài chính → có thể có thiên lệch nhớ lại (recall bias).', color=RGBColor(0xCC, 0, 0))
P(doc, '• Không phân tích đa biến (logistic regression với các biến cùng lúc) — chỉ phân tích '
       'từng cặp biến.', color=RGBColor(0xCC, 0, 0))

P(doc, 'Áp dụng cho đề tài VTN:', bold=True)
P(doc, 'Mặc dù đối tượng người trưởng thành, bài này cung cấp 2 bài học quý: '
       '(1) Hiện tượng PERFECTIONISM TRAP — học sinh/sinh viên giỏi nguy cơ cao, phù hợp với phát '
       'hiện của Trần Thị Mỵ Lương 2020 trên HS THPT chuyên (67,5 % cảm thấy "vô dụng"); '
       '(2) Vai trò của giấc ngủ — phù hợp Zheng 2025 (β = 0,615, mạnh nhất) và Zhu 2025 (AOR = '
       '13,71 cho < 5h). Đề cương VN cần đo CẢ 3 trục DASS + thêm câu hỏi giấc ngủ chi tiết.')

P(doc, 'Đánh giá: ⭐⭐⭐ Trung bình–Khá. Tạp chí TCNCYH trong nước, mẫu vừa, công cụ chuẩn, '
       'phát hiện perfectionism thú vị nhưng đối tượng không phù hợp trực tiếp đề tài VTN.',
  bold=True)

save(doc, os.path.join(DICH_DIR, '59_DaoThiNgoan_2025_TCNCYH_SVY4_HMU.docx'))

# 59.2 — Tóm tắt VN28
print('  Tóm tắt VN28...')
doc = make_doc()
H(doc, 'Tóm tắt bài VN28 — Thực trạng tâm lý sinh viên năm 4 ĐH Y Hà Nội', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Thực trạng tâm lý của sinh viên năm thứ tư Trường Đại học Y Hà Nội năm học '
       '2023-2024" của Đào Thị Ngoãn và cộng sự (2025), đăng trên Tạp chí Nghiên cứu Y học '
       '(TCNCYH) tập 187 số 02-2025, tr. 296-303. Khách thể: 196 sinh viên năm thứ 4 ĐH Y Hà '
       'Nội (126 ngành bác sĩ + 70 cử nhân, 35,71 % nam, 64,29 % nữ), được chọn ngẫu nhiên đơn '
       'từ danh sách 901 sinh viên toàn khối. Thời gian khảo sát: 01–12/2024. Công cụ: DASS-21 '
       'phiên bản tiếng Việt (Viện SKTT BV Bạch Mai chuẩn hoá) + Bộ câu hỏi nhu cầu tham vấn tâm '
       'lý do Viện ĐTYHDP&YTCC – ĐH Y Hà Nội phát triển năm 2019.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Thiết kế cắt ngang trên 196 sinh viên năm 4. Cỡ mẫu tính theo công thức ước tính tỷ lệ '
       'quần thể với p = 0,85, d = 0,05, α = 0,05 → n_min = 195. Sinh viên trả lời trực tiếp '
       'tại phòng LAB Bộ môn Sinh lý học. Nhập liệu Excel, phân tích STATA 15.0. Thống kê mô '
       'tả + χ²/Fisher + tính OR cho các liên quan có ý nghĩa thống kê (p < 0,05).')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ rối loạn sức khoẻ tâm thần ở SV năm 4 ĐH Y Hà Nội rất cao: 33,67 % trầm cảm, '
       '43,88 % LO ÂU, 31,63 % căng thẳng. 57,65 % mắc ít nhất 1 trong 3 rối loạn. 11,73 % '
       'mắc rối loạn từ mức nghiêm trọng trở lên (5,69 % mức nghiêm trọng + 6,12 % rất nghiêm '
       'trọng cộng dồn 3 trục).')

table(doc, ['Outcome', 'Yếu tố nguy cơ chính', 'OR', 'p'],
      [['Trầm cảm', 'Điểm tích luỹ Giỏi', '4,97', '0,005'],
       ['Lo âu', 'Điểm tích luỹ Giỏi', '4,69', '0,013'],
       ['Lo âu', 'Hệ Bác sĩ vs Cử nhân', '1,86', '0,044'],
       ['Lo âu', 'Vấn đề giấc ngủ', '1,84', '0,038'],
       ['Căng thẳng', 'Chưa hài lòng học tập', '10,70', '0,005'],
       ['Căng thẳng', 'Mối quan hệ xã hội có vấn đề', '2,92', '0,002'],
       ['Căng thẳng', 'Vấn đề giấc ngủ', '2,90', '0,001'],
       ['Căng thẳng', 'Tài chính không đủ', '2,74', '0,012']],
      widths=[3.0, 7.0, 2.0, 3.0])

P(doc, 'PHÁT HIỆN ĐẶC BIỆT — "Perfectionism trap": Sinh viên đạt điểm tích luỹ MỨC GIỎI có nguy '
       'cơ trầm cảm gấp 4,97 lần, lo âu gấp 4,69 lần, căng thẳng gấp 3,82 lần so với nhóm Khá – '
       'Trung bình. Đây là hiện tượng NGHỊCH ĐẢO so với giả thuyết thông thường, phù hợp với '
       'phát hiện của Trần Thị Mỵ Lương 2020 (HS THPT chuyên, 67,5 % cảm thấy "vô dụng").', bold=True)

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: chọn mẫu ngẫu nhiên đơn (không thuận tiện); DASS-21 chuẩn hoá tiếng Việt; '
       'đo cả 3 trục; phát hiện perfectionism trap thú vị. Hạn chế: đối tượng SINH VIÊN ĐẠI HỌC '
       '— không phải VTN học sinh THCS/THPT, không ngoại suy trực tiếp được; cỡ mẫu n = 196 nhỏ; '
       'cắt ngang không kết luận nhân quả; chỉ 1 trường ĐH Y Hà Nội; thiếu phân tích đa biến.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Lặp lại với HS THPT chuyên VN — dùng cùng công cụ DASS-Y (phiên bản VTN) để kiểm tra '
       'perfectionism trap ở tuổi nhỏ hơn. So sánh đa cơ sở (ĐH Y Hà Nội, ĐH Y Dược TPHCM, ĐH '
       'Y Huế). Phân tích mediation: điểm cao → áp lực → giấc ngủ → trầm cảm/lo âu. NC dọc theo '
       'dõi sinh viên Y năm 1 → 6 để xem điểm uốn.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐ Trung bình–Khá. TCNCYH trong nước, mẫu vừa, có giá trị bổ sung dữ liệu SV y VN.',
  bold=True)
save(doc, os.path.join(TT_DIR, 'VN28_DaoThiNgoan_2025_SVY4_HMU.docx'))

# ============================================================
# BÀI 60 — Duong et al. 2025 — Bản dịch tóm lược (full document)
# ============================================================
print('\n[60] Bài 60 — Duong 2025 Soc Psychiatry — bản dịch + tóm tắt')

doc = make_doc()
P(doc, 'Link: https://doi.org/10.1007/s00127-025-03043-7 | Social Psychiatry and Psychiatric '
       'Epidemiology (2025)', size=10)
H(doc, 'Phơi bày gánh nặng triệu chứng sức khoẻ tâm thần và hành vi nguy cơ ở vị thành niên Việt Nam: '
       'Bằng chứng từ nghiên cứu cắt ngang đa trung tâm trên 2.631 học sinh THPT', level=1)
p = doc.add_paragraph()
r = p.add_run('Unmasking the burden of mental health symptoms and risk behaviors in Vietnamese '
              'adolescents: evidence from a multicenter cross-sectional study involving 2,631 high '
              'school students')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

H(doc, 'THÔNG TIN THƯ MỤC', 2)
table(doc,
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Truc Thanh Thai, Hong-Tuyet Vo Le, Trang Thi Nguyen, Ngon Van Dinh, Xuan Le Mai, Hoai-Thuong Thi Tran, Ngoc-Bich Thi Nguyen, Khanh-Ha Mai Huynh, Thu-An Thi Nguyen, Hy-Han Thi Bui, Minh Cuong Duong*'],
        ['Cơ quan', '(1) Khoa YTCC, ĐH Y Dược TPHCM, 217 Hồng Bàng, Q.5, TPHCM; (2) School of Population Health, UNSW Sydney, Úc'],
        ['Tác giả liên hệ', 'Minh Cuong Duong — minh.duong@unsw.edu.au; Truc Thanh Thai — thaithanhtruc@ump.edu.vn'],
        ['Tạp chí', 'Social Psychiatry and Psychiatric Epidemiology (Q1, IF ≈ 4,5) — Springer'],
        ['Năm', '2025/2026 (online © The Author(s) 2026)'],
        ['DOI', '10.1007/s00127-025-03043-7'],
        ['Loại NC', 'Cắt ngang đa trung tâm (multicenter cross-sectional)'],
        ['Mẫu', '3.025 học sinh phát phiếu → 2.631 hợp lệ phân tích'],
        ['Địa bàn', '4 trường THPT + 4 Trung tâm Giáo dục thường xuyên (GDTX) tại TPHCM'],
        ['Công cụ', 'DASS-21 (sàng lọc trầm cảm/lo âu/căng thẳng) + YBRS (Youth Behavior Risk Scale)'],
        ['Truy cập', 'Springer (kiểm tra Open Access)'],
    ],
    widths=[3.5, 12.0])

H(doc, 'TÓM TẮT', 2)
P(doc, 'Bối cảnh: Vị thành niên thường gặp triệu chứng rối loạn tâm thần (SOMD) và tham gia vào '
       'các hành vi nguy cơ sức khoẻ (HRB) — cả hai đều là gánh nặng lớn về tử vong và tàn tật '
       'toàn cầu. Tuy nhiên, dữ liệu về vấn đề này còn hạn chế ở các nước thu nhập thấp – trung '
       'bình (LMIC), bao gồm Việt Nam. NC này nhằm ước tính tỷ lệ SOMD và HRB và xem xét mối '
       'quan hệ giữa chúng ở học sinh THPT Việt Nam.')

P(doc, 'Phương pháp: Khảo sát cắt ngang trên 3.025 học sinh từ 4 trường THPT và 4 trung tâm '
       'giáo dục thường xuyên (GDTX) tại TPHCM. Người tham gia hoàn thành bộ câu hỏi tự điền '
       'gồm thông tin nhân khẩu, HRB (sử dụng thang YBRS) và SOMD (sử dụng thang sàng lọc '
       'DASS-21). Đánh giá SOMD tập trung vào triệu chứng trầm cảm, lo âu, căng thẳng (không '
       'phải chẩn đoán lâm sàng). HRB bao phủ: sử dụng chất, hành vi nguy cơ, đánh nhau, ý '
       'nghĩ tự tử, hành vi tình dục không an toàn, ăn uống không lành mạnh, ít vận động, '
       'thiếu ngủ.')

P(doc, 'Kết quả: 2.631 học sinh được đưa vào phân tích cuối cùng. Tỷ lệ triệu chứng:')
table(doc, ['Triệu chứng', 'Tỷ lệ', 'Ghi chú'],
      [['Trầm cảm', '42,6 %', 'Cao'],
       ['LO ÂU', '50,3 %', 'Cao nhất'],
       ['Căng thẳng', '31,1 %', '—'],
       ['Hành vi nguy cơ thấp nhất', '4,0 %', 'Quan hệ tình dục không an toàn'],
       ['Hành vi nguy cơ cao nhất', '79,9 %', 'Ít vận động thể chất'],
       ['Đa hành vi nguy cơ (≥2 HRB)', '91,6 %', 'GẦN TUYỆT ĐỐI']],
      widths=[6.0, 3.0, 6.5])

P(doc, 'Học sinh có SOMD có xác suất tham gia HRB cao hơn đáng kể so với không có SOMD, với '
       'odds ratio dao động 1,24 đến 4,64.')

P(doc, 'Kết luận: SOMD và HRB là 2 thách thức KÉP và LIÊN KẾT ở vị thành niên Việt Nam — được '
       'nhấn mạnh bởi tỷ lệ rất cao. Phát hiện này nhấn mạnh nhu cầu cấp thiết cho các can thiệp '
       'TÍCH HỢP giải quyết đồng thời triệu chứng SKTT và hành vi nguy cơ, đặc biệt ở các bối '
       'cảnh LMIC nguồn lực hạn chế.')

H(doc, 'BỐI CẢNH NGHIÊN CỨU', 2)
P(doc, 'Việt Nam có 14 % dân số 97 triệu là vị thành niên. TPHCM là một trong 2 thành phố kinh '
       'tế-văn hoá-giáo dục hàng đầu, gồm 16 quận nội thành, 5 huyện ngoại thành và 1 thành phố '
       'vệ tinh (Thủ Đức). Hệ thống giáo dục Việt Nam chia làm nhiều cấp; THPT bao gồm lớp 10–12 '
       'với độ tuổi 15–18.')

P(doc, 'Đa số NC trước đây tại VN tập trung vào MỘT loại HRB hoặc một loại SOMD — nhưng trong '
       'thực tế các hành vi này LIÊN KẾT và đồng tồn tại. Theo lý thuyết hành vi vấn đề của '
       'Jessor, khi VTN tham gia nhiều hành vi nguy cơ là vì các hành vi đó liên kết và bị ảnh '
       'hưởng bởi môi trường. Tại Malaysia, > 80 % VTN đồng thời tham gia ≥ 2 HRB. Tại Ghana '
       '(11–19 tuổi), 94,8 % tham gia nhiều HRB. Tại các nước thu nhập cao (Mỹ, UK, Hà Lan), '
       '50–70 % tham gia ≥ 2 HRB.')

P(doc, 'Tại Việt Nam, một số sáng kiến hỗ trợ VTN đã được giới thiệu: Chương trình SK Học '
       'đường quốc gia, mở rộng dịch vụ tham vấn tâm lý học đường, các thí điểm can thiệp '
       'thúc đẩy SKTT trường học. Tuy nhiên, triển khai không đồng đều giữa các vùng, và '
       'thành phần SKTT thường chưa được tích hợp đầy đủ vào trường học.')

H(doc, 'PHÁT HIỆN CHÍNH', 2)

H(doc, 'Tỷ lệ SOMD (DASS-21) — n = 2.631', 3)
P(doc, '• Trầm cảm 42,6 % | Lo âu 50,3 % | Căng thẳng 31,1 %')
P(doc, '• Lo âu là triệu chứng phổ biến nhất — vượt cả trầm cảm.')
P(doc, '• So với các nghiên cứu khác tại VN: phù hợp với Hoa 2024 Hà Nội (40,6 % GAD-7), Long '
       'An 2025 (57,2 %), Hải Phòng 2024 (39,3 %).')

H(doc, 'Tỷ lệ HRB (YBRS)', 3)
P(doc, '• Ít vận động thể chất: 79,9 % (CAO NHẤT)')
P(doc, '• Đa hành vi nguy cơ (≥ 2 HRB): 91,6 % (gần như tuyệt đối)')
P(doc, '• Quan hệ tình dục không an toàn: 4,0 % (THẤP NHẤT)')
P(doc, '• Các HRB khác (sử dụng chất, đánh nhau, ý nghĩ tự tử, ăn uống không lành mạnh, '
       'thiếu ngủ): dao động trong khoảng giữa.')

H(doc, 'Mối liên hệ SOMD ↔ HRB', 3)
P(doc, 'Học sinh có triệu chứng SOMD có nguy cơ tham gia HRB cao hơn 1,24 đến 4,64 lần so với '
       'học sinh không có SOMD. Cụ thể (theo abstract):')
P(doc, '• Trầm cảm + lo âu + căng thẳng đều liên quan đến hầu hết các HRB.')
P(doc, '• OR cao nhất 4,64 — mối liên hệ rất mạnh.')
P(doc, '• OR thấp nhất 1,24 — mối liên hệ yếu hơn nhưng vẫn có ý nghĩa.')

H(doc, 'QUAN ĐIỂM PHẢN BIỆN', 2)
P(doc, 'Điểm mạnh:', bold=True, color=RGBColor(0xCC, 0, 0))
P(doc, '• Tạp chí Q1 — Social Psychiatry and Psychiatric Epidemiology (Springer, IF ≈ 4,5).', color=RGBColor(0xCC, 0, 0))
P(doc, '• Cỡ mẫu RẤT LỚN — n = 2.631 sau làm sạch (3.025 phát phiếu) — thuộc nhóm lớn nhất tại '
       'VN cho NC SKTT VTN.', color=RGBColor(0xCC, 0, 0))
P(doc, '• ĐA TRUNG TÂM — 4 trường THPT + 4 trung tâm GDTX — bao phủ cả học sinh chính quy lẫn '
       'giáo dục thường xuyên (đối tượng dễ bị bỏ qua).', color=RGBColor(0xCC, 0, 0))
P(doc, '• ĐO ĐỒNG THỜI SOMD và HRB — lấp khoảng trống lớn trong nghiên cứu VN (đa số NC trước '
       'chỉ đo 1 loại).', color=RGBColor(0xCC, 0, 0))
P(doc, '• Sử dụng DASS-21 + YBRS — 2 công cụ tin cậy quốc tế.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Khung lý thuyết Jessor (Problem Behavior Theory) + self-medication hypothesis — phân '
       'tích cơ chế.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Tỷ lệ lo âu 50,3 % — nhất quán với các NC VN khác (40–57 %), củng cố "baseline lo '
       'âu VTN châu Á" 40–55 %.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Đa hành vi nguy cơ 91,6 % — phát hiện CHẤN ĐỘNG, gợi ý cần can thiệp tích hợp đa '
       'thành phần.', color=RGBColor(0xCC, 0, 0))

P(doc, 'Hạn chế:', bold=True, color=RGBColor(0xCC, 0, 0))
P(doc, '• Cắt ngang — không thể kết luận chiều nhân quả (SOMD → HRB hay HRB → SOMD?).', color=RGBColor(0xCC, 0, 0))
P(doc, '• Chỉ TPHCM — không đại diện toàn VN. Không bao gồm các vùng nông thôn, miền núi, '
       'dân tộc thiểu số.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Tự báo cáo HRB → có thể có thiên lệch xã hội (social desirability bias) — đặc biệt '
       'với câu hỏi nhạy cảm như sử dụng chất, quan hệ tình dục, ý nghĩ tự tử.', color=RGBColor(0xCC, 0, 0))
P(doc, '• DASS-21 là sàng lọc, không phải chẩn đoán lâm sàng → có thể overestimate.', color=RGBColor(0xCC, 0, 0))
P(doc, '• Mất 13 % mẫu (3.025 → 2.631) — cần kiểm tra missing data analysis.', color=RGBColor(0xCC, 0, 0))

P(doc, 'Áp dụng cho đề tài VN:', bold=True)
P(doc, '• Đây là một trong các NC VN GIÁ TRỊ NHẤT về SOMD VTN — Q1 + cỡ mẫu lớn + đa trung tâm. '
       'Có thể trích dẫn làm nguồn baseline chính cho TPHCM.')
P(doc, '• Phát hiện 91,6 % đa hành vi nguy cơ → can thiệp đề cương VN cần TÍCH HỢP nhiều mục '
       'tiêu (CBT + giáo dục sức khoẻ + thay đổi hành vi + hỗ trợ gia đình), không chỉ '
       'targeted vào lo âu đơn lẻ.')
P(doc, '• OR 1,24–4,64 SOMD ↔ HRB → cơ chế tự dùng (self-medication) — phù hợp với phát '
       'hiện Zheng 2025 (MXH gián tiếp 63 % qua self-efficacy) và Dong 2025 (kênh giao '
       'tiếp gia đình OR = 0,22).')
P(doc, '• Nên lặp lại tại Hà Nội + miền Bắc/Trung để có bức tranh toàn quốc.')

P(doc, 'Đánh giá: ⭐⭐⭐⭐⭐ Rất cao. Q1, mẫu lớn, đa trung tâm, đo đồng thời SOMD + HRB — bài VN '
       'tốt nhất hiện có cho dữ liệu VTN TPHCM hiện đại.', bold=True)

save(doc, os.path.join(DICH_DIR, '60_Duong_2025_SocPsychiatry_2631HS_TPHCM.docx'))

# 60.2 — Tóm tắt VN29
print('  Tóm tắt VN29...')
doc = make_doc()
H(doc, 'Tóm tắt bài VN29 — Triệu chứng SKTT và hành vi nguy cơ ở 2.631 HS THPT TPHCM', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Unmasking the burden of mental health symptoms and risk behaviors in '
       'Vietnamese adolescents: evidence from a multicenter cross-sectional study involving '
       '2,631 high school students" của Truc Thanh Thai, Minh Cuong Duong và cộng sự (2025), '
       'đăng trên Social Psychiatry and Psychiatric Epidemiology (Q1, IF ≈ 4,5) — Springer. '
       'DOI 10.1007/s00127-025-03043-7. Khách thể: 2.631 học sinh THPT (sau làm sạch từ 3.025) '
       'tại 4 trường THPT + 4 Trung tâm Giáo dục thường xuyên TPHCM. Đơn vị: Khoa YTCC, ĐH Y Dược '
       'TPHCM (217 Hồng Bàng, Q.5) hợp tác với UNSW Sydney, Úc. Công cụ: DASS-21 (SOMD) + YBRS '
       '(Youth Behavior Risk Scale, HRB).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Khảo sát cắt ngang đa trung tâm tại 8 cơ sở giáo dục TPHCM (4 trường THPT chính quy + '
       '4 trung tâm giáo dục thường xuyên). Học sinh tự điền bộ câu hỏi gồm 3 phần: (1) Nhân '
       'khẩu xã hội; (2) DASS-21 sàng lọc trầm cảm/lo âu/căng thẳng; (3) YBRS đo 8 nhóm hành '
       'vi nguy cơ (sử dụng chất, hành vi nguy hiểm, đánh nhau, ý nghĩ tự tử, hành vi tình dục, '
       'ăn uống, vận động, giấc ngủ). Phân tích logistic regression để đánh giá mối liên hệ giữa '
       'SOMD và HRB. Khung lý thuyết: Jessor Problem Behavior Theory + self-medication hypothesis.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ triệu chứng (DASS-21):')
table(doc, ['Triệu chứng', 'Tỷ lệ'],
      [['Trầm cảm', '42,6 %'],
       ['Lo âu (cao nhất)', '50,3 %'],
       ['Căng thẳng', '31,1 %']],
      widths=[8.0, 5.5])

P(doc, 'Tỷ lệ hành vi nguy cơ (HRB):')
table(doc, ['HRB', 'Tỷ lệ'],
      [['Ít vận động thể chất (cao nhất)', '79,9 %'],
       ['Quan hệ tình dục không an toàn (thấp nhất)', '4,0 %'],
       ['Đa hành vi nguy cơ (≥ 2 HRB)', '91,6 %']],
      widths=[10.0, 3.5])

P(doc, 'Mối liên hệ: Học sinh có SOMD có nguy cơ tham gia HRB cao hơn 1,24 đến 4,64 lần so với '
       'học sinh không có SOMD — bằng chứng mạnh cho cơ chế "tự dùng" (self-medication) ở VTN '
       'Việt Nam. Phù hợp với khung lý thuyết Jessor về liên kết các hành vi vấn đề.', bold=True)

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Q1 Social Psychiatry and Psychiatric Epidemiology, cỡ mẫu rất lớn (n = 2.631), '
       'đa trung tâm bao phủ cả THPT chính quy + GDTX (đối tượng dễ bị bỏ qua), đo ĐỒNG THỜI '
       'SOMD và HRB — lấp khoảng trống lớn trong NC VN. Sử dụng khung lý thuyết quốc tế. Tỷ lệ '
       'lo âu 50,3 % phù hợp với "baseline VTN châu Á 40–55 %" — củng cố bức tranh dịch tễ.')
P(doc, 'Hạn chế: cắt ngang không kết luận chiều nhân quả; chỉ TPHCM (không đại diện cả nước); '
       'tự báo cáo có thể có thiên lệch xã hội với câu hỏi nhạy cảm; DASS-21 sàng lọc có thể '
       'overestimate; mất 13 % mẫu cần kiểm tra missing data analysis.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Lặp lại tại Hà Nội + miền Trung để có dữ liệu toàn quốc. NC dọc theo dõi nhóm này 2-3 '
       'năm để xác định chiều nhân quả SOMD ↔ HRB. Phát triển và thử nghiệm CAN THIỆP TÍCH HỢP '
       'đa thành phần (CBT + giáo dục SK + thay đổi hành vi + hỗ trợ gia đình) thay vì targeted '
       'vào lo âu/trầm cảm đơn lẻ — vì 91,6 % HS đa hành vi nguy cơ. So sánh THPT chính quy vs '
       'GDTX về cơ cấu nguy cơ.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao. Q1, mẫu lớn, đa trung tâm, đo đồng thời SOMD + HRB — bài VN tốt nhất hiện có cho dữ liệu VTN TPHCM.',
  bold=True)
save(doc, os.path.join(TT_DIR, 'VN29_Duong_2025_2631HS_TPHCM.docx'))

print('\n=== DONE ===')
print('Bản dịch đầy đủ:')
print(f'  - 03_Ban-dich/59_DaoThiNgoan_2025_TCNCYH_SVY4_HMU.docx')
print(f'  - 03_Ban-dich/60_Duong_2025_SocPsychiatry_2631HS_TPHCM.docx')
print('Tóm tắt:')
print(f'  - Tom-tat-tung-bai/VN28_DaoThiNgoan_2025_SVY4_HMU.docx')
print(f'  - Tom-tat-tung-bai/VN29_Duong_2025_2631HS_TPHCM.docx')
