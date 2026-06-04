# -*- coding: utf-8 -*-
"""Audit broad kho - kiem tra cac bai trong tong quan LA con co loi khong.
27/05/2026.

Check:
- Cac trich dan trong LA co khop voi ban dich + tom tat khong
- Author name spelling consistency
- Number consistency
"""
import os, sys, io, re, glob
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Cac bai duoc trich dan trong LA (theo TLTK)
# Format: (cite_in_la, n_value_expected, prevalence_expected, file_paper)
KEY_PAPERS = [
    ('Saikia', '288', '24,4%', 'QT002_Saikia_et_al_2023'),
    ('Mandaknalli', None, None, 'QT003_Mandaknalli_Malusare_2021'),
    ('Chen', None, '13,9%', 'QT007_Chen_et_al_2023_China'),
    ('Wen', None, None, 'QT008_Wen_et_al_2020_China'),
    ('Qiu', None, None, 'QT009_Qiu_et_al_2022_China'),
    ('Xu', '373.216', None, 'QT010_Xu_et_al_2021_China'),
    ('Bhardwaj', None, '81,9%', 'QT011_Bhardwaj_et_al_2020'),
    ('Zhameden', None, None, 'QT013_Zhameden_2025'),
    ('Anderson', None, None, 'QT014_Anderson_2025'),
    ('Zhu', None, None, 'QT015_Zhu_2025_BMC_China'),
    ('Mudunna', None, None, 'QT016_Mudunna_2025'),
    ('Puyat', None, None, 'QT017_Puyat_2025_Filipino'),
    ('Salari', None, None, 'QT018_Salari_SAD'),
    ('Shibuya', None, None, 'QT019_Shibuya_SchoolMH'),
]

# Load LA v3_2
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_2_FixSoLieu_27052026.docx')
d = Document(LA)
la_full_text = '\n'.join(p.text for p in d.paragraphs)

print("=" * 70)
print("AUDIT - LA v3_2 vs kho ban-dich")
print("=" * 70)
print(f"LA paragraphs: {len(d.paragraphs)}, total chars: {len(la_full_text)}")
print()

for author_key, n_expected, prev_expected, file_glob in KEY_PAPERS:
    print(f"\n--- {author_key} ({file_glob}) ---")
    # Find in LA
    matches_in_la = []
    for i, p in enumerate(d.paragraphs):
        if author_key in p.text and i < 1300:  # Only body, not TLTK
            matches_in_la.append(i)
    print(f"  LA mentions: {len(matches_in_la)} paras (body only)")
    if matches_in_la:
        for idx in matches_in_la[:3]:
            txt = d.paragraphs[idx].text
            # Try to extract numbers near author name
            # Find sentence with author
            pos = txt.find(author_key)
            ctx = txt[max(0, pos-80):pos+200]
            print(f"    para {idx}: ...{ctx}...")
    # Check ban-dich exists
    bd_pattern = os.path.join(ROOT, '03_Ban-dich', f'*{file_glob}*.docx')
    bd_files = [f for f in glob.glob(bd_pattern) if '_FIXED_' not in f and '_backup' not in f.lower()]
    print(f"  Ban-dich files: {[os.path.basename(f) for f in bd_files]}")
    # Check tom-tat exists
    tt_pattern = os.path.join(ROOT, 'Tom-tat-tung-bai', f'*{file_glob}*.docx')
    tt_files = [f for f in glob.glob(tt_pattern) if '_FIXED_' not in f]
    print(f"  Tom-tat files: {[os.path.basename(f) for f in tt_files]}")
