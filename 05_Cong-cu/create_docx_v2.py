"""
Script v2: Create .docx with PDF page images inserted INLINE
at correct positions (after corresponding sections).
Tables/figures appear right where they're referenced.
"""
import os
import fitz  # PyMuPDF
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
import glob

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
TRANS_DIR = os.path.join(BASE, "Translations")
ORIG_DIR = os.path.join(BASE, "Originals")
OUT_DIR = os.path.join(BASE, "DocFiles")
IMG_DIR = os.path.join(BASE, "PDF_Images")

os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(IMG_DIR, exist_ok=True)

# Mapping: which PDF pages to insert after which sections
# Format: {paper_key: [(section_keyword, [page_numbers]), ...]}
# Pages are 1-indexed
PAGE_MAP = {
    "01_Jenkins_et_al_2023": [
        ("TÓM TẮT", [1]),           # page 1: title + abstract
        ("PHƯƠNG PHÁP", [2, 3]),     # pages 2-3: methods + recruitment
        ("KẾT QUẢ", [4, 5, 6]),     # pages 4-6: Table 1, Table 2, Figure 1, Figure 2
        ("THẢO LUẬN", [7, 8, 9, 10, 11]),  # pages 7-11: discussion + refs
    ],
    "02_Saikia_et_al_2023": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [2]),
        ("KẾT QUẢ", [3, 4, 5]),     # Table 1, 2, 3
        ("THẢO LUẬN", [5, 6]),
    ],
    "03_Mandaknalli_Malusare_2021": [
        ("TÓM TẮT", [1]),
        ("KẾT QUẢ", [2, 3]),
        ("THẢO LUẬN", [3, 4]),
    ],
    "05_Alharbi_et_al_2019": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [2]),
        ("KẾT QUẢ", [2, 3, 4, 5]),  # Figure 1, 2, Table 1, 2, 3
        ("THẢO LUẬN", [5, 6, 7]),
    ],
    "06_Nakie_et_al_2022": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [2, 3, 4]),
        ("KẾT QUẢ", [4, 5, 6, 7, 8, 9]),  # Tables, figures
        ("THẢO LUẬN", [9, 10, 11, 12]),
    ],
    "07_Chen_et_al_2023": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [2, 3]),
        ("KẾT QUẢ", [3, 4, 5, 6]),
        ("THẢO LUẬN", [6, 7, 8, 9]),
    ],
    "08_Wen_et_al_2020": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [2, 3, 4, 5]),
        ("KẾT QUẢ", [5, 6, 7, 8, 9, 10]),
        ("THẢO LUẬN", [10, 11, 12, 13, 14]),
    ],
    "09_Qiu_et_al_2022": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [2, 3]),
        ("KẾT QUẢ", [3, 4, 5, 6]),   # Figure 1, Tables 1-3
        ("THẢO LUẬN", [6, 7, 8, 9]),
    ],
    "10_Xu_et_al_2021": [
        ("TÓM TẮT", [1]),
        ("PHƯƠNG PHÁP", [1, 2]),
        ("KẾT QUẢ", [2, 3, 4, 5]),   # Table 1, Fig 1-3, Table 2-3
        ("THẢO LUẬN", [5, 6]),
    ],
    "11_Bhardwaj_et_al_2020": [
        ("TÓM TẮT", [1, 2]),
        ("PHƯƠNG PHÁP", [3, 4, 5]),
        ("KẾT QUẢ", [5, 6, 7, 8, 9, 10]),
        ("THẢO LUẬN", [10, 11, 12, 13, 14, 15]),
    ],
}


def extract_pdf_pages(pdf_path, prefix):
    """Extract pages as images, return dict {page_num: img_path}."""
    pages = {}
    try:
        doc = fitz.open(pdf_path)
        for i, page in enumerate(doc):
            mat = fitz.Matrix(2.0, 2.0)
            pix = page.get_pixmap(matrix=mat)
            img_path = os.path.join(IMG_DIR, f"{prefix}_p{i+1}.png")
            pix.save(img_path)
            pages[i + 1] = img_path
        doc.close()
    except Exception as e:
        print(f"  Error: {e}")
    return pages


def insert_image(doc, img_path, caption=""):
    """Insert an image into the document."""
    try:
        doc.add_picture(img_path, width=Inches(6.3))
        last_p = doc.paragraphs[-1]
        last_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if caption:
            p = doc.add_paragraph()
            run = p.add_run(caption)
            run.font.size = Pt(9)
            run.font.italic = True
            run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    except Exception as e:
        doc.add_paragraph(f"[Lỗi chèn hình: {e}]")


def add_word_table(doc, table_lines, in_critical=False):
    """Parse markdown table lines and create a real Word table."""
    # Filter out separator lines (|---|---|)
    data_lines = []
    for tl in table_lines:
        tl = tl.strip()
        if tl and not all(c in '|-: ' for c in tl):
            cells = [c.strip() for c in tl.strip('|').split('|')]
            data_lines.append(cells)

    if not data_lines:
        return

    num_cols = max(len(row) for row in data_lines)
    # Pad rows to same column count
    for row in data_lines:
        while len(row) < num_cols:
            row.append('')

    table = doc.add_table(rows=len(data_lines), cols=num_cols)
    table.style = 'Table Grid'

    for i, row_data in enumerate(data_lines):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            # Handle bold markers
            clean = cell_text.strip()
            if '**' in clean:
                parts = clean.split('**')
                for k, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        if k % 2 == 1:
                            run.bold = True
                        if in_critical:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(clean)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                if in_critical:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

            # Bold header row
            if i == 0:
                for run in p.runs:
                    run.bold = True

    doc.add_paragraph()  # spacing after table


def md_to_docx_v2(md_path, docx_path, pdf_pages=None, page_map=None):
    """Convert markdown to docx with inline PDF images and real Word tables."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_critical_review = False
    inserted_pages = set()
    current_section = ""
    table_buffer = []  # Collect table lines

    for idx, line in enumerate(lines):
        line = line.rstrip('\n')

        # Check if we're in a table
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            table_buffer.append(line)
            continue
        elif table_buffer:
            # Table just ended — render it
            add_word_table(doc, table_buffer, in_critical_review)
            table_buffer = []

        if not line.strip():
            continue

        if 'QUAN ĐIỂM PHẢN BIỆN' in line or 'CRITICAL REVIEW' in line:
            in_critical_review = True

        # Detect section changes and insert images BEFORE new section starts
        if line.startswith('## ') and pdf_pages and page_map:
            # Check if previous section had images to insert
            if current_section:
                for kw, page_nums in page_map:
                    if kw in current_section:
                        for pn in page_nums:
                            if pn in pdf_pages and pn not in inserted_pages:
                                insert_image(doc, pdf_pages[pn],
                                           f"[Trang {pn} bản gốc / Original page {pn}]")
                                inserted_pages.add(pn)
                        break

            current_section = line

        # Render markdown
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            if in_critical_review:
                for run in h.runs:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('---'):
            pass  # skip horizontal rules
        elif line.startswith('- **') or line.startswith('  - '):
            p = doc.add_paragraph(style='List Bullet')
            text = line.lstrip(' -')
            if '**' in text:
                parts = text.split('**')
                for j, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                        if j % 2 == 1:
                            run.bold = True
                        if in_critical_review:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                if in_critical_review:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('- '):
            p = doc.add_paragraph(line[2:], style='List Bullet')
            if in_critical_review:
                for run in p.runs:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        else:
            p = doc.add_paragraph()
            text = line
            if '**' in text:
                parts = text.split('**')
                for j, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(12)
                        if j % 2 == 1:
                            run.bold = True
                        if in_critical_review:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(text)
                run.font.size = Pt(12)
                if in_critical_review:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    # Flush remaining table buffer
    if table_buffer:
        add_word_table(doc, table_buffer, in_critical_review)
        table_buffer = []

    # Insert any remaining pages after last section
    if current_section and pdf_pages and page_map:
        for kw, page_nums in page_map:
            if kw in current_section:
                for pn in page_nums:
                    if pn in pdf_pages and pn not in inserted_pages:
                        insert_image(doc, pdf_pages[pn],
                                   f"[Trang {pn} bản gốc / Original page {pn}]")
                        inserted_pages.add(pn)
                break

    # Insert any pages not yet inserted at the very end
    if pdf_pages:
        remaining = [pn for pn in sorted(pdf_pages.keys()) if pn not in inserted_pages]
        if remaining:
            doc.add_page_break()
            doc.add_heading("CÁC TRANG GỐC BỔ SUNG / ADDITIONAL ORIGINAL PAGES", level=2)
            for pn in remaining:
                insert_image(doc, pdf_pages[pn],
                           f"[Trang {pn} bản gốc / Original page {pn}]")

    doc.save(docx_path)
    print(f"  ✅ {os.path.basename(docx_path)}")


# PDF map
pdf_map = {
    "01_Jenkins_et_al_2023": "01_Jenkins_et_al_2023.pdf",
    "02_Saikia_et_al_2023": "02_Saikia_et_al_2023.pdf",
    "03_Mandaknalli_Malusare_2021": "03_Mandaknalli_Malusare_2021.pdf",
    "04_NSCH_2020": None,
    "05_Alharbi_et_al_2019": "05_Alharbi_et_al_2019.pdf",
    "06_Nakie_et_al_2022": "06_Nakie_et_al_2022.pdf",
    "07_Chen_et_al_2023": "07_Chen_et_al_2023.pdf",
    "08_Wen_et_al_2020": "08_Wen_et_al_2020.pdf",
    "09_Qiu_et_al_2022": "09_Qiu_et_al_2022.pdf",
    "10_Xu_et_al_2021": "10_Xu_et_al_2021.pdf",
    "11_Bhardwaj_et_al_2020": "11_Bhardwaj_et_al_2020.pdf",
}

print("=" * 60)
print("TẠO DOCX V2 — HÌNH ẢNH CHÈN ĐÚNG VỊ TRÍ")
print("=" * 60)

# Extract images
print("\n--- Trích xuất hình ảnh PDF ---")
all_pages = {}
for key, pdf_name in pdf_map.items():
    if pdf_name:
        pdf_path = os.path.join(ORIG_DIR, pdf_name)
        if os.path.exists(pdf_path):
            pages = extract_pdf_pages(pdf_path, key)
            all_pages[key] = pages
            print(f"  {pdf_name}: {len(pages)} trang")

# Create docx
print("\n--- Tạo file DOCX ---")
for md_file in sorted(glob.glob(os.path.join(TRANS_DIR, "*.md"))):
    basename = os.path.splitext(os.path.basename(md_file))[0]
    docx_path = os.path.join(OUT_DIR, f"{basename}.docx")
    pages = all_pages.get(basename)
    pmap = PAGE_MAP.get(basename)
    md_to_docx_v2(md_file, docx_path, pages, pmap)

# Summary docx
print("\n--- Tạo bảng tóm tắt ---")
summary_md = os.path.join(BASE, "BANG_TOM_TAT_11_BAI_BAO.md")
summary_docx = os.path.join(OUT_DIR, "BANG_TOM_TAT_TONG_HOP.docx")
md_to_docx_v2(summary_md, summary_docx)

print(f"\n{'='*60}")
print("HOÀN TẤT!")
print(f"{'='*60}")
