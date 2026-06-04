# -*- coding: utf-8 -*-
"""Dịch + tóm tắt 5 bài QT mới"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

papers = [
    # (fname_dich, fname_tomtat, title_vn, title_en, journal, n_pages, info_rows, results_headers, results_rows, results_widths, findings, critique, gap, rating)
    (
        '30_GBD_Trends_10-24y_2025.docx',
        'QT30_GBD_Trends.docx',
        'Xu hướng rối loạn trầm cảm và lo âu ở VTN và thanh niên (10\u201324 tuổi) từ 1990 đến 2021: Phân tích GBD',
        'Trends in Depressive and Anxiety Disorders among Adolescents and Young Adults (Aged 10\u201324) from 1990 to 2021: A GBD Study Analysis',
        'Journal of Affective Disorders, 2025 (Q1, IF\u22486,6). Zhang D et al. 14 trang.',
        [('Phạm vi', '204 nước/vùng lãnh thổ, 1990\u20132021 (GBD 2021)'),
         ('Đối tượng', 'VTN và thanh niên 10\u201324 tuổi'),
         ('Phương pháp', 'Joinpoint regression + Age-Period-Cohort analysis'),
         ('Chỉ số', 'Tỷ lệ, tỷ suất mới mắc, DALYs cho trầm cảm + lo âu')],
        ['Chỉ số', 'Toàn cầu', 'Xu hướng'],
        [['Lo âu 10\u201314 tuổi', '3,6%', 'Tăng, đặc biệt 2014\u20132021'],
         ['Lo âu 15\u201319 tuổi', '4,6%', 'Tăng'],
         ['Trầm cảm 10\u201314', '2,1%', 'Tăng'],
         ['Trầm cảm 15\u201319', '2,8%', 'Tăng'],
         ['AAPC tăng nhanh nhất', '2014\u20132021', 'Đỉnh 2019'],
         ['SDI cao', 'Tỷ lệ cao nhất', 'Nước thu nhập cao gánh nặng lớn']],
        [3.5, 3.5, 4.5],
        'Gánh nặng lo âu+trầm cảm tăng TOÀN CẦU, đặc biệt 2014\u20132021 và nước SDI cao. Joinpoint xác định điểm bẻ 2014.',
        'Q1 IF=6,6, GBD 2021, 204 nước. Hạn chế: dữ liệu GBD phụ thuộc báo cáo quốc gia (VN thiếu).',
        'Phân tích riêng ASEAN/Việt Nam từ GBD. So sánh xu hướng VN vs nước SDI cao.',
        '\u2b50\u2b50\u2b50\u2b50\u2b50 Q1, GBD 2021, 204 nước, joinpoint nâng cao.',
    ),
    (
        '31_59Countries_Anxiety_2025.docx',
        'QT31_59Countries.docx',
        'Tỷ lệ và các yếu tố liên quan của lo âu ở VTN đi học: Phân tích từ 59 quốc gia',
        'Prevalence of and Factors Associated with Anxiety among School Going Adolescents: Analysis from 59 Countries',
        'Journal of Affective Disorders, 2025 (Q1, IF\u22486,6). Islam MA et al. 11 trang.',
        [('Phạm vi', '59 quốc gia, dữ liệu GSHS (Global School-based Student Health Survey)'),
         ('Đối tượng', 'VTN đi học, nhiều vùng WHO'),
         ('Phương pháp', 'Phân tích đa quốc gia, hồi quy đa cấp')],
        ['Vùng WHO', 'Tỷ lệ lo âu', 'Ghi chú'],
        [['Đông Địa Trung Hải', '14,6%', 'Cao nhất'],
         ['Châu Phi', '11,0%', 'Thứ 2'],
         ['Châu Mỹ', '~8\u201310%', '\u2014'],
         ['Tây Thái Bình Dương', '~5\u20138%', '\u2014'],
         ['Đông Nam Á', '3,6%', 'THẤP NHẤT'],
         ['Yếu tố nguy cơ', 'Bắt nạt, cô đơn, thiếu bạn thân', 'Liên quốc gia'],
         ['Yếu tố bảo vệ', 'Hỗ trợ cha mẹ, giáo dục cha mẹ', 'Liên quốc gia']],
        [3.5, 3.0, 4.5],
        'ĐNA có tỷ lệ lo âu THẤP NHẤT (3,6%). Tuy nhiên, dữ liệu GSHS có thể đánh giá thấp do phương pháp.',
        'Q1, 59 nước, GSHS chuẩn hóa. Hạn chế: GSHS chỉ 1\u20132 câu hỏi lo âu (không phải GAD-7).',
        'So sánh GSHS vs GAD-7 trên cùng mẫu VTN ĐNA. Giải thích tại sao ĐNA thấp nhất.',
        '\u2b50\u2b50\u2b50\u2b50 Q1, 59 nước, so sánh liên quốc gia.',
    ),
    (
        '32_Ireland_MyWorld_2024.docx',
        'QT32_Ireland.docx',
        'Khám phá xu hướng thay đổi trầm cảm và lo âu ở VTN từ 2012 đến 2019: Từ khảo sát My World',
        'Exploring Changing Trends in Depression and Anxiety among Adolescents from 2012 to 2019: Insights from My World Repeated Cross-sectional Surveys',
        'Early Intervention in Psychiatry, 2024. Fitzgerald A et al. Ireland. 11 trang.',
        [('Phạm vi', 'Ireland, 2 đợt khảo sát 2012 và 2019'),
         ('Đối tượng', 'VTN Ireland'),
         ('Phương pháp', 'Cắt ngang lặp lại (repeated cross-sectional)')],
        ['Chỉ số', 'Phát hiện'],
        [['Xu hướng 2012\u21922019', 'Trầm cảm và lo âu TĂNG có ý nghĩa, đặc biệt ở NỮ'],
         ['Yếu tố dự báo', 'Giải thích 37\u201361% phương sai'],
         ['Yếu tố nguy cơ', 'Bắt nạt, mạng xã hội, stress'],
         ['Yếu tố bảo vệ', 'Lòng tự trọng, hỗ trợ xã hội, hoạt động thể chất'],
         ['Giới tính', 'Nữ tăng nhanh hơn nam'],
         ['Khuyến nghị', 'Ưu tiên dịch vụ SKTT cho nữ VTN']],
        [4.0, 8.0],
        'Xu hướng tăng 7 năm Ireland, nữ tăng nhanh hơn. Phù hợp Na Uy (2025) và JAACAP (2024).',
        'Thiết kế cắt ngang lặp lại — tốt hơn đơn cắt ngang. Hạn chế: chỉ Ireland, tự báo cáo.',
        'So sánh xu hướng Ireland vs VN. Can thiệp nhắm nữ VTN.',
        '\u2b50\u2b50\u2b50\u2b50 Cắt ngang lặp lại, xu hướng 7 năm, yếu tố giải thích.',
    ),
    (
        '33_JAMA_ScreenMedia_2024.docx',
        'QT33_JAMA_Screen.docx',
        'Sử dụng phương tiện màn hình và sức khỏe tâm thần trẻ em/VTN: Phân tích thứ cấp RCT',
        'Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial',
        'JAMA Network Open, 2024. Schmidt-Persson J et al. Denmark. 12 trang.',
        [('Phạm vi', 'Đan Mạch, phân tích thứ cấp RCT'),
         ('Đối tượng', 'Trẻ em và VTN'),
         ('Phương pháp', 'Phân tích thứ cấp từ thử nghiệm lâm sàng ngẫu nhiên')],
        ['Chỉ số', 'Phát hiện'],
        [['Thiết kế', 'Phân tích thứ cấp RCT — bằng chứng mạnh hơn quan sát'],
         ['Screen time', 'Liên quan với triệu chứng SKTT ở trẻ em/VTN'],
         ['Cơ chế', 'Giảm thời gian màn hình cải thiện SKTT'],
         ['So sánh', 'JAMA — tạp chí hàng đầu thế giới (IF \u2248 13,0)']],
        [4.0, 8.0],
        'Bằng chứng từ RCT (không chỉ quan sát): giảm screen time cải thiện SKTT. JAMA IF=13.',
        'JAMA Q1 IF=13, RCT — bằng chứng nhân quả mạnh. Hạn chế: phân tích thứ cấp, chỉ Đan Mạch.',
        'RCT giảm screen time tại trường VN. So sánh bối cảnh Đan Mạch vs châu Á.',
        '\u2b50\u2b50\u2b50\u2b50\u2b50 JAMA Q1 IF=13, bằng chứng RCT.',
    ),
    (
        '34_Korea_MH_Trends_2024.docx',
        'QT34_Korea_Trends.docx',
        'Xu hướng quốc gia về SKTT VTN theo thu nhập tại Hàn Quốc, trước và sau COVID-19, 2006\u20132022',
        'National Trends in Adolescents\u2019 Mental Health by Income Level in South Korea, Pre\u2013 and Post\u2013COVID-19, 2006\u20132022',
        'Nature Scientific Reports, 2024 (Q1, IF\u22484,6). Hàn Quốc. 10 trang. KYRBS 2006\u20132022.',
        [('Phạm vi', 'Hàn Quốc, KYRBS quốc gia 2006\u20132022 (16 năm)'),
         ('Đối tượng', 'VTN THCS + THPT'),
         ('Phương pháp', 'Phân tích xu hướng theo thu nhập, COVID')],
        ['Chỉ số', '2006/trước COVID', '2022/sau COVID', 'Xu hướng'],
        [['Stress (thu nhập cao)', '\u2014', '40,1%', 'Giảm trước COVID, tăng sau'],
         ['Stress (thu nhập thấp)', '\u2014', '62,8%', 'Luôn cao hơn'],
         ['Buồn bã', '\u2014', '28,2%', 'Tăng sau COVID'],
         ['Ý tưởng tự tử', '\u2014', '13,9%', 'Tăng sau COVID'],
         ['Cố tự tử', '\u2014', '\u2014', 'Tăng'],
         ['Bất bình đẳng thu nhập', 'Chênh lệch', 'Mở rộng sau COVID', 'Thu nhập thấp tệ hơn']],
        [4.0, 3.0, 3.0, 3.0],
        'Xu hướng 16 năm: SKTT cải thiện trước COVID, xấu đi sau COVID. Thu nhập thấp luôn tệ hơn + khoảng cách mở rộng.',
        'Nature Q1, KYRBS quốc gia, 16 năm. Hạn chế: chỉ Hàn Quốc, đo stress/buồn (không tách lo âu riêng).',
        'So sánh xu hướng Hàn Quốc vs VN. Phân tích bất bình đẳng thu nhập trong SKTT VTN VN.',
        '\u2b50\u2b50\u2b50\u2b50 Nature Q1, KYRBS 16 năm, bất bình đẳng thu nhập.',
    ),
]

for fname_dich, fname_tomtat, title_vn, title_en, journal, info_rows, res_h, res_rows, res_widths, findings, critique, gap, rating in papers:
    # === BẢN DỊCH ===
    d = create_doc()
    add_heading(d, title_vn, 1)
    add_heading(d, title_en, 2)
    add_heading(d, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(d, [('Tạp chí', journal)] + list(info_rows))

    add_heading(d, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    add_p(d, f'\u2022 {findings}', bold=True)
    add_p(d, f'\u2022 Phản biện: {critique}')
    add_p(d, f'\u2022 Gap: {gap}')

    add_heading(d, 'KẾT QUẢ CHÍNH', 2)
    add_table(d, res_h, res_rows, res_widths)

    add_heading(d, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW', 2)
    add_red(d, critique)
    add_red(d, f'Khoảng trống: {gap}')

    add_p(d, f'Đánh giá: {rating}', bold=True)
    d.save(fname_dich)
    sys.stderr.write(f'{fname_dich} OK\n')

    # === TÓM TẮT ===
    from docx import Document
    from docx.shared import Pt, RGBColor
    RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

    d2 = Document()
    s = d2.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    from docx.shared import Cm
    for sec in d2.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)

    p = d2.add_paragraph(); r = p.add_run(f'Tóm tắt: {title_vn[:80]}...'); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE

    p2 = d2.add_paragraph(); r2 = p2.add_run('Tên công trình'); r2.bold = True; r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.color.rgb = RED
    p3 = d2.add_paragraph(); r3 = p3.add_run(f'{title_vn}. {journal}'); r3.font.name = 'Times New Roman'; r3.font.size = Pt(12); r3.font.color.rgb = BLUE

    p4 = d2.add_paragraph(); r4 = p4.add_run('Kết quả chính'); r4.bold = True; r4.font.name = 'Times New Roman'; r4.font.size = Pt(12); r4.font.color.rgb = RED

    # Add table to summary too
    from docx.enum.table import WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    def shade2(cell, color):
        s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); cell._tc.get_or_add_tcPr().append(s)
    def set_w2(cell, w):
        tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)

    t = d2.add_table(rows=1+len(res_rows), cols=len(res_h)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(res_h)): set_w2(row.cells[ci], res_widths[ci])
    for i, h in enumerate(res_h):
        c = t.rows[0].cells[i]; c.text = h
        for pp in c.paragraphs:
            pp.alignment = 1
            for rr in pp.runs: rr.bold = True; rr.font.name = 'Times New Roman'; rr.font.size = Pt(10)
        shade2(c, 'D9E2F3')
    for ri, rd in enumerate(res_rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for pp in c.paragraphs:
                for rr in pp.runs: rr.font.name = 'Times New Roman'; rr.font.size = Pt(10)

    h5 = d2.add_heading('Phản biện', 2)
    for r in h5.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
    p5 = d2.add_paragraph(); r5 = p5.add_run(critique); r5.font.name = 'Times New Roman'; r5.font.size = Pt(12); r5.font.color.rgb = RED

    p6 = d2.add_paragraph(); r6 = p6.add_run(f'Đánh giá: {rating}'); r6.bold = True; r6.font.name = 'Times New Roman'; r6.font.size = Pt(12)

    d2.save(os.path.join('..', 'Tom-tat-tung-bai', fname_tomtat))
    sys.stderr.write(f'{fname_tomtat} OK\n')
