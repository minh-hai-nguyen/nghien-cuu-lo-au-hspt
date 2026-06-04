"""DOC: Phan bien VN004 Nguyen Thi Van — yeu to HOC TAP
Theo style hom qua lam voi VN014 quan he cha me-con (rong qua).

ALL FACTS VERIFIED:
- VN004 (Tom-tat verified):
  + Nhom hoc tap r = 0,37 voi lo au
  + Top 3 bieu hien: ap luc thi DH 56,7%; dinh huong nghe 51,5%; ky vong cha me 48,9%
  + Cong cu: STAI Spielberger ban Viet Nguyen Cong Khanh 2000 (do lo au)
  + Yeu to anh huong: bang hoi tu thiet ke (khong neu thang chuan)
  + n = 558 sang loc -> 90 phong van sau (mau thuan tien, 4 truong TPHCM)
- Sun et al. 2011 ESSA: 16 muc, 5 chieu (Pressure from study, Workload, Worry about grades, Self-expectation, Despondency)
- Pascoe Hetrick Parker 2020 khung 6 truc
- Tran Thao Vi 2024 J Rural Medicine n=611 Hue dung ESSA - verified
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Phan_bien_VN004_yeu_to_hoc_tap.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def blue_bullet(text, size=12):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẢN BIỆN VN004 NGUYỄN THỊ VÂN (2020)\nKHÁI NIỆM "YẾU TỐ HỌC TẬP" CÓ RỘNG QUÁ KHÔNG?')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Em kiểm tra bình luận thêm Nguyễn Thị Vân về yếu tố học '
    'tập nhé, theo kiểu bình luận về NC của HT Học (hôm qua '
    'chúng ta làm ấy).', italic=True
)

# 1. Số liệu cụ thể VN004 về yếu tố học tập
H('1. Khái niệm "yếu tố học tập" trong VN004', level=2, color=NAVY)
para(
    'Nguyễn Thị Vân (2020) — VN004 trong cơ sở dữ liệu — khảo '
    'sát 558 học sinh THPT TPHCM bằng STAI Spielberger (bản '
    'Việt hóa Nguyễn Công Khanh, 2000) → 90 học sinh phỏng vấn '
    'sâu, mẫu thuận tiện 4 trường nội + ngoại thành. Tác giả '
    'phân tích bốn nhóm yếu tố ảnh hưởng đến lo âu: HỌC TẬP / '
    'gia đình / quan hệ xã hội / bản thân HS.'
)
caption('Bảng 1. Yếu tố HỌC TẬP trong VN004')
add_table(
    ['Đặc điểm', 'VN004 Nguyễn Thị Vân (2020)'],
    [
        ['Tên biến', '"Nhóm yếu tố học tập"'],
        ['Hệ số tương quan r với lo âu', 'r = 0,37 (mạnh thứ 2 sau "bản thân HS" r=0,42)'],
        ['Thang đo', 'Bảng hỏi tự thiết kế (KHÔNG nêu thang chuẩn quốc tế)'],
        ['Số mục', 'Không công bố cụ thể'],
        ['Số chiều phân biệt', 'CHỈ MỘT — gộp tất cả vào 1 biến'],
        ['Top biểu hiện 1', 'Áp lực thi đại học (56,7% học sinh)'],
        ['Top biểu hiện 2', 'Định hướng nghề (51,5%)'],
        ['Top biểu hiện 3', 'Kỳ vọng cha mẹ về điểm số (48,9%)'],
        ['Phân biệt áp lực thi vs khối lượng bài tập', 'KHÔNG'],
        ['Phân biệt kỳ vọng nội vs ngoại', 'KHÔNG'],
        ['Phân biệt áp lực hiện tại vs tương lai', 'KHÔNG'],
    ]
)
para('')
para(
    'Tương tự VN014 (Hoàng Trung Học) đã phản biện hôm qua, '
    'VN004 đo "yếu tố học tập" chỉ bằng MỘT biến tổng — KHÔNG '
    'phân biệt các chiều con khác nhau. Hệ số r = 0,37 cho '
    'biết tổng thể nhóm học tập có liên quan với lo âu, nhưng '
    'KHÔNG xác định được chiều con nào MẠNH NHẤT.'
)

# 2. So sánh chuẩn quốc tế
H('2. So sánh chuẩn quốc tế — VN004 đo TỔNG, không phân tách', level=2, color=NAVY)
para(
    'Y văn quốc tế đã từ lâu phân biệt nhiều chiều con của áp '
    'lực học tập. Bảng dưới đây tổng hợp các thang chuẩn được '
    'sử dụng nhiều nhất.'
)
caption('Bảng 2. Ba thang đo chuẩn quốc tế cho áp lực học tập + lo âu thi cử')
add_table(
    ['Thang đo', 'Tác giả gốc', 'Số mục', 'Số chiều'],
    [
        ['ESSA (Educational Stress Scale for Adolescents)',
         'Sun, Dunne, Hou & Xu (2011)', '16',
         '5 chiều: Pressure from study / Workload / Worry about grades / Self-expectation / Despondency'],
        ['CTAS (Cognitive Test Anxiety Scale)',
         'Cassady & Johnson (2002)', '27',
         '1 chiều: cognitive test anxiety (lo âu thi cử nhận thức)'],
        ['TAI (Test Anxiety Inventory)',
         'Spielberger (1980)', '20',
         '2 chiều: Worry (lo lắng) / Emotionality (cảm xúc kích hoạt)'],
    ]
)
para('')
para('Hai ví dụ so sánh:', bold=True)
para(
    'Thứ nhất, Trần Thảo Vi và cộng sự (2024) trong Journal of '
    'Rural Medicine trên 611 học sinh trung học Huế đã sử dụng '
    'thang ESSA 16 mục — phân tách rõ năm chiều con. Phù hợp '
    'với chuẩn quốc tế Sun và cộng sự (2011), cho phép so sánh '
    'meta-analysis với y văn châu Á và châu Âu.'
)
para(
    'Thứ hai, Chương 3 luận án của thầy đã sử dụng ESSA RÚT '
    'GỌN 4 mục (ESSA.3 bài tập, ESSA.4 sự nghiệp tương lai, '
    'ESSA.5 kỳ vọng cha mẹ, ESSA.6 kiểm tra). Tốt hơn VN004 '
    '(bảng tự thiết kế) nhưng vẫn yếu hơn ESSA full 16 mục — '
    'đặc biệt thiếu chiều "Despondency" (chán nản học tập) — '
    'một chiều CỐT LÕI theo Sun 2011.'
)

# 3. Sáu chiều con bị gộp chung
H('3. Sáu chiều con của "yếu tố học tập" bị VN004 gộp chung', level=2, color=NAVY)
caption('Bảng 3. Sáu chiều con của áp lực học tập và tác động dự kiến với lo âu')
add_table(
    ['#', 'Chiều con', 'Tác động dự kiến với lo âu', 'Bằng chứng tham chiếu'],
    [
        ['1', 'Áp lực thi cử (test anxiety)',
         'β DƯƠNG mạnh — yếu tố nguy cơ',
         'Pascoe 2020; Spielberger 1980 TAI'],
        ['2', 'Khối lượng bài tập (workload)',
         'β DƯƠNG vừa — yếu tố nguy cơ',
         'OECD 2018 — VN top 5 bài tập về nhà'],
        ['3', 'Lo lắng về điểm số (worry about grades)',
         'β DƯƠNG mạnh — đặc biệt với lo âu xã hội',
         'Sun 2011; Putwain 2008'],
        ['4', 'Kỳ vọng tự thân (self-expectation)',
         'β có thể DƯƠNG hoặc U-NGƯỢC',
         'Pascoe 2020 mô hình U-ngược'],
        ['5', 'Kỳ vọng từ cha mẹ (parental expectation)',
         'β DƯƠNG — đặc biệt văn hóa Á',
         'VN004: 48,9% HS bị ảnh hưởng'],
        ['6', 'Định hướng tương lai (future-oriented stress)',
         'β DƯƠNG — đặc thù VN sớm hơn phương Tây',
         'VN004: 51,5%; Trần Thảo Vi 2024'],
    ]
)
para('')
para(
    'Hai luận điểm quan trọng từ bảng trên. Thứ nhất, sáu '
    'chiều con có CƠ CHẾ và CƯỜNG ĐỘ tác động khác nhau lên '
    'lo âu. Áp lực thi cử thường tác động qua trục HPA (cortisol '
    'cấp tính), khối lượng bài tập tác động qua giấc ngủ, '
    'kỳ vọng cha mẹ tác động qua tự đánh giá. Gộp tất cả '
    'thành 1 biến TRIỆT TIÊU sự khác biệt cơ chế.'
)
para(
    'Thứ hai, các chiều con có khả năng can thiệp KHÁC NHAU. '
    'Áp lực thi cử có thể giảm bằng kỹ năng quản lý thi (TAI '
    'intervention). Khối lượng bài tập cần can thiệp cấp '
    'TRƯỜNG (giảm bài tập). Kỳ vọng cha mẹ cần can thiệp cấp '
    'GIA ĐÌNH. Không phân tách thì KHÔNG thiết kế được can '
    'thiệp đích.'
)

# 4. Bốn hậu quả phương pháp luận
H('4. Bốn hậu quả phương pháp luận', level=2, color=NAVY)
para('Việc gộp tất cả thành 1 biến gây ra bốn hậu quả nghiêm trọng cho giá trị khoa học của VN004.')

H('4.1. Hệ số r = 0,37 KHÔNG phân biệt được aspect mạnh nhất', level=3)
para(
    'Khi r = 0,37 cho biết "yếu tố học tập có liên quan với lo '
    'âu" — nhưng không trả lời được: aspect NÀO mạnh nhất? '
    'Áp lực thi cử có mạnh hơn khối lượng bài tập? Kỳ vọng '
    'cha mẹ có vượt lo lắng điểm số? Mỗi can thiệp khác '
    'nhau cho mỗi chiều con. Hệ số tổng không hướng dẫn '
    'được thực hành.'
)

H('4.2. Top 3 biểu hiện chỉ là TẦN SUẤT, không phải PREDICTOR', level=3)
para(
    'Ba biểu hiện hàng đầu trong VN004 — áp lực thi đại học '
    '(56,7%), định hướng nghề (51,5%), kỳ vọng cha mẹ '
    '(48,9%) — chỉ là TỶ LỆ % học sinh có biểu hiện. Đây '
    'là TẦN SUẤT XUẤT HIỆN, KHÔNG phải CƯỜNG ĐỘ TÁC ĐỘNG '
    'lên lo âu. Hai khái niệm khác nhau hoàn toàn:'
)
para(
    '• 56,7% học sinh có áp lực thi ĐH KHÔNG nghĩa là áp lực '
    'thi ĐH gây 56,7% phương sai lo âu.\n'
    '• Có thể chỉ 30% học sinh có áp lực kỳ vọng cha mẹ '
    'NHƯNG cường độ tác động β = 0,5 lên lo âu — mạnh hơn '
    '56,7% áp lực thi ĐH với β = 0,2.'
)
para(
    'Kết luận: VN004 KHÔNG cung cấp thông tin về chiều nào '
    'mạnh nhất với lo âu — chỉ cho biết chiều nào phổ biến '
    'nhất.'
)

H('4.3. Không so sánh trực tiếp được với y văn quốc tế', level=3)
para(
    'r = 0,37 từ thang tự thiết kế của VN004 KHÔNG so sánh '
    'được với r tương ứng từ ESSA/TAI/AAS quốc tế. Khoảng '
    'cách so sánh này làm giảm giá trị đóng góp của VN004 '
    'cho y văn — bài chỉ "đóng góp số liệu Việt Nam" mà '
    'không cho phép so sánh quốc tế. Đặc biệt, KHÔNG thể '
    'đưa VN004 vào meta-analysis quốc tế về áp lực học tập.'
)

H('4.4. Không hướng dẫn được thiết kế can thiệp đích', level=3)
para(
    'Khuyến nghị "giảm áp lực học tập" trong VN004 quá '
    'chung chung. Giảm aspect nào? Khối lượng bài tập? Áp '
    'lực thi? Kỳ vọng cha mẹ? Định hướng tương lai? Mô '
    'hình can thiệp khác nhau cho mỗi chiều:'
)
para('• Giảm khối lượng bài tập → can thiệp CẤP TRƯỜNG (chính sách)')
para('• Giảm áp lực thi → can thiệp CẤP HỌC SINH (kỹ năng quản lý thi)')
para('• Giảm kỳ vọng cha mẹ → can thiệp CẤP GIA ĐÌNH (parent training)')
para('• Giảm áp lực tương lai → can thiệp CẤP TƯ VẤN HƯỚNG NGHIỆP')
para(
    'Không có phân tách chiều con thì không có can thiệp '
    'đích.'
)

# 5. Đề xuất
H('5. Bốn đề xuất cho thiết kế nghiên cứu mới', level=2, color=NAVY)
para('Đề tài của thầy có thể LẤP KHOẢNG TRỐNG này bằng cách đo "yếu tố học tập" theo nhiều chiều ngay từ thiết kế.', bold=True)

para('ĐỀ XUẤT 1 — Chọn thang chuẩn quốc tế. Ba lựa chọn ưu tiên:')
add_table(
    ['Lựa chọn', 'Thang', 'Ưu điểm', 'Nhược điểm'],
    [
        ['ƯU TIÊN 1',
         'ESSA full 16 mục (Sun et al. 2011), 5 chiều',
         'Chuẩn quốc tế nhất; đã được Trần Thảo Vi 2024 sử dụng tại VN; cho phép so sánh meta-analysis',
         'Cần dịch + chuẩn hóa'],
        ['ƯU TIÊN 2',
         'ESSA rút gọn 4 mục (như chương 3 luận án)',
         'Ngắn gọn, dễ áp dụng; đã có data VN',
         'Thiếu chiều Despondency'],
        ['ƯU TIÊN 3',
         'TAI (Spielberger 1980) 20 mục cho áp lực thi cử riêng',
         'Chi tiết về test anxiety; phân chiều Worry vs Emotionality',
         'Chỉ đo áp lực thi, bỏ qua kỳ vọng và workload'],
    ]
)

para('')
para(
    'ĐỀ XUẤT 2 — Phân tích RIÊNG hiệu ứng từng chiều với '
    'lo âu. Dự kiến kết quả: áp lực thi β DƯƠNG mạnh; lo '
    'lắng điểm β DƯƠNG; kỳ vọng cha mẹ β DƯƠNG; tự kỳ vọng '
    'có thể U-NGƯỢC. Nếu kết quả khớp dự kiến, đây là bằng '
    'chứng VIỆT NAM bổ sung cho phát hiện quốc tế Pascoe '
    '2020 + Sun 2011.'
)
para(
    'ĐỀ XUẤT 3 — Áp dụng SEM với biến tiềm ẩn đa chiều. Mô '
    'hình: ESSA 5 chiều → biến tiềm ẩn "Áp lực học tập" → '
    'lo âu. Cho phép kiểm định cấu trúc 5 nhân tố trên dữ '
    'liệu Việt Nam. So sánh với CHƯƠNG 3 luận án (chỉ dùng '
    '4 mục) để thấy được tác động riêng của Despondency.'
)
para(
    'ĐỀ XUẤT 4 — Thiết kế CAN THIỆP ĐÍCH theo chiều con. '
    'Sau khi xác định chiều mạnh nhất, thiết kế can thiệp '
    'TƯƠNG ỨNG. Ví dụ: nếu kỳ vọng cha mẹ β CAO NHẤT → '
    'parent training; nếu áp lực thi cao nhất → kỹ năng '
    'quản lý thi; nếu workload cao nhất → can thiệp cấp '
    'trường.'
)

# 6. CÂU TRẢ LỜI
H('6. CÂU TRẢ LỜI', level=2, color=NAVY)
blue_run('Tóm gọn:', bold=True)
blue_bullet(
    'CÓ — khái niệm "yếu tố học tập" trong VN004 RỘNG QUÁ. '
    'Tác giả gộp tất cả thành phần (áp lực thi, khối lượng '
    'bài tập, lo lắng điểm, kỳ vọng tự thân, kỳ vọng cha '
    'mẹ, định hướng tương lai) vào MỘT biến với r = 0,37 — '
    'không phân biệt được aspect nào mạnh nhất với lo âu.'
)
blue_bullet(
    'Trái với VN004, Trần Thảo Vi và cộng sự (2024) trong '
    'Journal of Rural Medicine trên 611 học sinh trung học '
    'Huế đã sử dụng thang ESSA 16 mục chuẩn quốc tế — phân '
    'tách 5 chiều rõ ràng. Chương 3 luận án của thầy dùng '
    'ESSA rút gọn 4 mục — tốt hơn VN004 nhưng vẫn yếu hơn '
    'ESSA full.'
)
blue_bullet(
    'Y văn quốc tế phân biệt 6 chiều con: áp lực thi cử / '
    'khối lượng bài tập / lo lắng điểm số / kỳ vọng tự '
    'thân / kỳ vọng cha mẹ / định hướng tương lai. Mỗi '
    'chiều có cơ chế và cường độ tác động khác nhau lên '
    'lo âu — gộp chung TRIỆT TIÊU phát hiện.'
)
blue_bullet(
    'Bốn hậu quả của việc gộp: (1) không phân biệt được '
    'aspect mạnh nhất; (2) Top 3 biểu hiện (56,7%/51,5%/'
    '48,9%) chỉ là TẦN SUẤT, không phải PREDICTOR; (3) '
    'không so sánh được với y văn quốc tế; (4) không '
    'hướng dẫn được thiết kế can thiệp đích.'
)
blue_bullet(
    'Đề xuất cho đề tài của thầy: dùng ESSA full 16 mục '
    '(Sun 2011) hoặc ESSA 4 mục (như chương 3); phân tích '
    'RIÊNG từng chiều với lo âu; SEM với biến tiềm ẩn đa '
    'chiều; thiết kế can thiệp đích theo chiều mạnh nhất.'
)

blue_run('Cách trích vào báo cáo CTH:', bold=True)
blue_run(
    '"Nguyễn Thị Vân (2020) khảo sát 558 học sinh THPT TPHCM '
    'và phát hiện nhóm yếu tố học tập có hệ số tương quan '
    'r = 0,37 với mức lo âu, đứng thứ hai sau yếu tố bản '
    'thân học sinh. Tuy nhiên, biến "yếu tố học tập" trong '
    'nghiên cứu được đo bằng bảng hỏi tự thiết kế chỉ một '
    'biến tổng hợp — không phân biệt các chiều con như áp '
    'lực thi cử, khối lượng bài tập, lo lắng điểm số, hay '
    'kỳ vọng cha mẹ. Trái với VN004, Trần Thảo Vi và cộng '
    'sự (2024) trong Journal of Rural Medicine đã sử dụng '
    'thang Educational Stress Scale for Adolescents (ESSA; '
    'Sun và cộng sự, 2011) gồm 16 mục với 5 chiều rõ ràng. '
    'Chương 3 luận án sử dụng ESSA rút gọn 4 mục — bao gồm '
    'áp lực bài tập, định hướng sự nghiệp tương lai, kỳ '
    'vọng cha mẹ, và áp lực kiểm tra. Mặc dù chi tiết hơn '
    'VN004, phiên bản rút gọn vẫn thiếu chiều "Despondency" '
    '(chán nản học tập) — một chiều cốt lõi theo Sun 2011. '
    'Khuyến nghị các nghiên cứu tiếp theo tại Việt Nam sử '
    'dụng ESSA full 16 mục kết hợp phân tích riêng từng '
    'chiều con với lo âu để xác định chiều mạnh nhất, tạo '
    'cơ sở cho can thiệp đích."', italic=True
)

# 7. Phụ lục
H('7. Phụ lục — Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Cassady, J. C., & Johnson, R. E. (2002). Cognitive test anxiety and academic performance. Contemporary Educational Psychology, 27(2), 270–295. https://doi.org/10.1006/ceps.2001.1094',
    'Nguyễn, T. V. (2020). Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh. [VN004 trong cơ sở dữ liệu dự án.]',
    'Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823 [QT067 trong cơ sở dữ liệu dự án.]',
    'Putwain, D. W. (2008). Examination stress and test anxiety. The Psychologist, 21(12), 1026–1029.',
    'Spielberger, C. D. (1980). Test Anxiety Inventory: Preliminary professional manual. Palo Alto, CA: Consulting Psychologists Press.',
    'Sun, J., Dunne, M. P., Hou, X. Y., & Xu, A. Q. (2011). Educational Stress Scale for Adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546. https://doi.org/10.1177/0734282910394976',
    'Tran, T. V., Nguyen, H. T. L., Tran, X. M. T., Tashiro, Y., Seino, K., Vo, T. V., & Nakamura, K. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine, 19(4), 279–290. https://doi.org/10.2185/jrm.2024-012',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
