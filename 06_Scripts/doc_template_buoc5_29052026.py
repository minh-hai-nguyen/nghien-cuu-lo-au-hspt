# -*- coding: utf-8 -*-
"""Doc template Buoc 5 - Giay tiep nhan PBDL lan 1.
29/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TEMPLATE = os.path.join(ROOT, 'Luận án TS',
                         'Bước 5. Giấy tiếp nhận Hồ sơ xin ý kiến Phản biện độc lập (lần 1) (1).docx')

print(f'Path: {TEMPLATE}')
print(f'Exists: {os.path.exists(TEMPLATE)}')
print()

d = Document(TEMPLATE)

print('=== PARAGRAPHS ===')
for i, para in enumerate(d.paragraphs, 1):
    if para.text.strip():
        align = para.alignment
        # Get font info from first run if exists
        font_info = ''
        if para.runs:
            r = para.runs[0]
            font_info = f' [size={r.font.size.pt if r.font.size else "?"} bold={r.bold} italic={r.italic}]'
        print(f'[{i}]{font_info} | align={align} | {para.text}')

print()
print('=== TABLES ===')
for ti, tbl in enumerate(d.tables, 1):
    print(f'--- Table {ti} ({len(tbl.rows)} rows x {len(tbl.columns)} cols) ---')
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                print(f'  R{ri}C{ci}: {txt}')

print()
print('=== SECTIONS / MARGINS ===')
for si, sec in enumerate(d.sections, 1):
    print(f'Section {si}: W={sec.page_width.cm:.2f}cm H={sec.page_height.cm:.2f}cm')
    print(f'  Margins: T={sec.top_margin.cm:.2f} B={sec.bottom_margin.cm:.2f} '
          f'L={sec.left_margin.cm:.2f} R={sec.right_margin.cm:.2f}')

print()
print('=== STYLE NORMAL ===')
s = d.styles['Normal']
print(f'Font: {s.font.name}, Size: {s.font.size.pt if s.font.size else "?"}')
print(f'Line spacing: {s.paragraph_format.line_spacing}')
