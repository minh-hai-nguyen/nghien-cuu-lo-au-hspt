# -*- coding: utf-8 -*-
"""VN15 + VN16 — viết lại chuẩn CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def rb(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(doc, t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ============================================================
# VN15: NGÔ ANH VINH 2024 — DTTS Lạng Sơn
# ============================================================
d = make_doc()
bl(d, 'Tóm tắt bài VN-15', bold=True)

rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d, 'Công trình nghiên cứu \u00ab Sức khỏe tâm thần ở thanh thiếu niên dân tộc thiểu số tại Việt Nam và các yếu tố liên quan: Nghiên cứu cắt ngang \u00bb (Mental Health among Ethnic Minority Adolescents in Vietnam and Correlated Factors: A Cross-sectional Study), do Ngô Anh Vinh, Vũ Thị Mỹ Hạnh, Đỗ Thị Bích Vân, Dương Anh Tài, Đỗ Minh Loan và Lê Thị Thanh Thùy (2024), thuộc Khoa Sức khỏe Vị thành niên, Bệnh viện Nhi Trung ương, Hà Nội, khảo sát trên mẫu 845 học sinh dân tộc thiểu số nội trú tại tỉnh Lạng Sơn, Việt Nam, xuất bản trên Journal of Affective Disorders Reports, tập 17, bài 100795 (Elsevier, Open Access).')

rb(d, 'Phương pháp nghiên cứu')
bl(d, 'Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21 \u2014 Depression Anxiety Stress Scale) phiên bản tiếng Việt (Cronbach alpha: trầm cảm 0,767; lo âu 0,792; căng thẳng 0,811) kết hợp với Bảng hỏi Trải nghiệm Bất lợi Thời thơ ấu (ACEs \u2014 Adverse Childhood Experiences Questionnaire). Nói cách khác, nghiên cứu này kết hợp hai chiều đánh giá \u2014 tình trạng sức khỏe tâm thần hiện tại và trải nghiệm bất lợi quá khứ \u2014 để hiểu mối quan hệ giữa "vết thương thời thơ ấu" và rối loạn cảm xúc.', bold=True)

bl(d, 'DASS-21 là phiên bản rút gọn 21 mục của thang gốc 42 mục (Lovibond & Lovibond, 1995). Ngưỡng cắt theo Henry & Crawford (2005): trầm cảm \u22655, lo âu \u22654, căng thẳng \u22658.')
bl(d, 'ACEs gồm 10 loại trải nghiệm bất lợi: lạm dụng thể chất, cảm xúc, tình dục; bỏ mặc thể chất, cảm xúc; bạo lực gia đình; nghiện chất; bệnh tâm thần cha mẹ; ly hôn; tù đày.')

bl(d, 'Tổng quan tài liệu của chúng tôi cho thấy dữ liệu về sức khỏe tâm thần thanh thiếu niên dân tộc thiểu số nội trú vùng cao Việt Nam còn rất hạn chế. UNICEF (2017) báo cáo 68,4% trẻ 1\u201314 tuổi VN trải qua kỷ luật thể chất hoặc tâm lý. Nghiên cứu trước (Thai và cộng sự, 2020) trên 4.957 VTN cho thấy 86% có ít nhất 1 ACE. Đây là nghiên cứu đầu tiên nhắm đến nhóm DTTS nội trú.', bold=True)

bl(d, 'Dữ liệu bao gồm tuổi, giới tính, dân tộc (Tày/Nùng/khác), tần suất về nhà, tần suất cha mẹ thăm, số người sống cùng, chất lượng tình bạn, hoạt động thể chất, tần suất sử dụng internet, nguồn thông tin SKTT, chiến lược ứng phó.')
bl(d, 'Nghiên cứu cắt ngang tại trường nội trú DTTS tỉnh Lạng Sơn, n = 845. Phân tích: hồi quy Logistic đa biến (OR) cho biến nhị phân + hồi quy Tobit (Coef.) cho điểm liên tục.')

rb(d, 'Kết quả nghiên cứu định lượng')
bl(d, 'Tỷ lệ sức khỏe tâm thần và ACEs (N = 845):', bold=True)
tbl(d,
    ['Tình trạng', 'Tỷ lệ', 'Ghi chú'],
    [['Trầm cảm (DASS-21 \u22655)', '59,0%', 'Cao nhất trong 3 chỉ số'],
     ['Lo âu (DASS-21 \u22654)', '54,4%', '\u2014'],
     ['Căng thẳng (DASS-21 \u22658)', '24,7%', '\u2014'],
     ['ACEs (\u22651 trải nghiệm)', '48,9%', 'TB = 1,1; SD = 1,8; range 0\u20136']],
    widths=[4.0, 2.5, 5.0])
d.add_paragraph()

bl(d, 'Yếu tố liên quan (hồi quy Logistic + Tobit đa biến):', bold=True)
tbl(d,
    ['Yếu tố', 'Coef./OR', 'KTC 95%', 'Ý nghĩa'],
    [['Số ACEs \u2192 lo âu', 'Coef. = 0,28', '\u2014', 'Dương, có ý nghĩa'],
     ['Số ACEs \u2192 căng thẳng', 'Coef. = 0,28', '\u2014', 'Dương, có ý nghĩa'],
     ['Số ACEs \u2192 trầm cảm', 'Coef. = 0,16', '\u2014', 'Dương, có ý nghĩa'],
     ['Bạn bè kém \u2192 trầm cảm', 'OR = 6,84***', '2,03\u201323,02', 'RẤT MẠNH'],
     ['Bạn bè kém \u2192 lo âu', 'Coef. = 0,54\u20131,46', '\u2014', 'Fair/Poor đều tăng'],
     ['Sống \u226510 người \u2192 lo âu', 'OR = 0,51***', '0,33\u20130,77', 'Bảo vệ'],
     ['Hoạt động thể chất \u2192 lo âu', 'OR = 0,72**', '0,51\u20131,00', 'Bảo vệ'],
     ['Internet 5h/ngày \u2192 lo âu', 'OR = 0,43**', '0,21\u20130,86', 'Bảo vệ (vs cuối tuần)'],
     ['TT từ bạn bè \u2192 lo âu', 'OR = 0,57*', '\u2014', 'Bảo vệ']],
    widths=[4.0, 2.5, 2.5, 3.5])
d.add_paragraph()

bh3(d, 'Khác biệt giới tính : Tổng quan tài liệu cho thấy nữ giới thường có tỷ lệ lo âu cao hơn nam (Salk và cộng sự, 2017). Tuy nhiên, nghiên cứu này không phân tích giới tính chi tiết cho lo âu \u2014 chỉ có OR = 1,29 cho căng thẳng ở nữ, không đạt ý nghĩa (P > 0,05). Đây là khoảng trống quan trọng khi văn hóa giới ở vùng DTTS có thể khác biệt so với vùng Kinh.')

rb(d, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3(d, 'Phát hiện về ACEs, tình bạn và bối cảnh trường nội trú')
bl(d, '*ACEs và sức khỏe tâm thần.* Nổi bật nhất là mối liên hệ dương tính giữa số ACEs và tất cả 3 chỉ số SKTT (Coef. = 0,16\u20130,28). Gần 50% HS DTTS nội trú đã trải qua ít nhất 1 ACE \u2014 cao hơn trung bình các nước LMIC (Blum và cộng sự, 2019) nhưng thấp hơn NC trước tại VN (Thai 2020: 86%). Điều này có thể do HS nội trú ít về nhà, ít tiếp xúc ACEs trực tiếp.')
bl(d, '*Mối quan hệ bạn bè.* Bạn bè kém (Poor) tăng nguy cơ trầm cảm cực mạnh: OR = 6,84 (KTC 2,03\u201323,02). Trong bối cảnh trường nội trú \u2014 nơi HS sống xa gia đình \u2014 bạn bè trở thành nguồn hỗ trợ chính. Phát hiện này tương tự NSCH 2020 (khó kết bạn gấp 10 lần) và Nag và cộng sự (2019) tại Tripura, Ấn Độ.')
bl(d, '*Internet hàng ngày.* Sử dụng internet hàng ngày (5h) bảo vệ hơn cuối tuần (OR = 0,43) \u2014 có thể do sử dụng đều đặn giúp kết nối xã hội tốt hơn. Nguồn thông tin từ trường KHÔNG liên quan SKTT \u2014 gợi ý vai trò SKTT của trường chưa đủ mạnh.')

rb(d, 'Kết luận')
bl(d, 'Dữ liệu của chúng tôi, cho thấy tỷ lệ rất cao trầm cảm (59,0%), lo âu (54,4%) và ACEs (48,9%) ở thanh thiếu niên dân tộc thiểu số nội trú tại Lạng Sơn \u2014 nghiên cứu đầu tiên nhắm đến nhóm đặc biệt dễ bị tổn thương này tại Việt Nam, gợi ý rằng can thiệp cần nhắm vào: giảm ACEs, cải thiện mối quan hệ bạn bè (OR trầm cảm = 6,84 nếu bạn bè kém), và tăng cường dịch vụ SKTT tại trường nội trú. Trong bối cảnh trường nội trú DTTS vùng cao, nhu cầu phát triển chương trình hỗ trợ tâm lý nhạy cảm văn hóa là cấp thiết.', bold=True)

rh2(d, 'Phản biện')
bl(d, 'Nghiên cứu đầu tiên DTTS nội trú VN, kết hợp DASS-21 + ACEs, hồi quy Logistic + Tobit đa biến, Elsevier Open Access. Tuy nhiên, chỉ 1 tỉnh (Lạng Sơn) \u2014 không đại diện cho tất cả DTTS VN. DASS-21 ngưỡng thấp (lo âu \u22654) \u2014 so V-NAMHS 2022 (DISC-5: 2,3%), chênh lệch 24 lần. Không phân tích giới tính chi tiết. ACEs tự báo cáo hồi cứu \u2014 recall bias. Tạp chí J Affective Disorders Reports (không phải J Affective Disorders Q1) \u2014 chưa có IF ổn định.')

rh2(d, 'Hướng nghiên cứu tiếp theo')
bl(d, 'Mở rộng đến nhiều tỉnh vùng cao (Hà Giang, Lai Châu, Sơn La). So sánh DTTS nội trú vs không nội trú vs Kinh cùng vùng. Nghiên cứu dọc ACEs \u2192 SKTT. Can thiệp cải thiện tình bạn tại trường nội trú. Phân tích giới tính chi tiết ở DTTS.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50 Trung bình-Khá. NC đầu tiên DTTS nội trú VN, DASS-21+ACEs, hồi quy đa biến, Elsevier OA, nhưng 1 tỉnh và DASS ngưỡng thấp.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.save('VN15_NgoAnhVinh_2024.docx')
sys.stderr.write('VN15 OK\n')

# ============================================================
# VN16: BẢO QUYÊN 2025 — Hà Nội
# ============================================================
d2 = make_doc()
bl(d2, 'Tóm tắt bài VN-16', bold=True)

rb(d2, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d2, 'Công trình nghiên cứu \u00ab Thực trạng sức khỏe tâm thần của học sinh THPT tại Hà Nội: Nghiên cứu cắt ngang về trầm cảm, lo âu và căng thẳng \u00bb (Mental Health Status of High School Students in Hanoi: A Cross-Sectional Study on Depression, Anxiety, and Stress), do Nguyễn Ngọc Bảo Quyên, Lưu Minh Đức, Vũ Minh Thy, Nguyễn Yến Minh, Nguyễn Châu Anh, Hoàng Lan Vân và Giáp Thị Thanh Tình (2025), thuộc Đại học VinUni và Đại học Y Dược ĐHQG Hà Nội, khảo sát trên 501 học sinh THPT tại Hà Nội (78% nữ) từ tháng 8 đến 12/2023 bằng khảo sát trực tuyến, xuất bản trên Tạp chí Y học Cộng đồng, tập 66, chuyên đề 10, 2025, trang 79\u201386.')

rb(d2, 'Phương pháp nghiên cứu')
bl(d2, 'Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21) qua khảo sát trực tuyến. Nói cách khác, nghiên cứu sử dụng phương pháp sàng lọc cắt ngang tự báo cáo trực tuyến \u2014 với ngưỡng cắt bao gồm cả mức "nhẹ", cho tỷ lệ rất cao.', bold=True)

bl(d2, 'DASS-21 là phiên bản rút gọn của thang gốc 42 mục (Lovibond & Lovibond, 1995). Phân tích bằng SPSS 27.0: thống kê mô tả, ANOVA, t-test, chi-square, hồi quy đa biến.')

bl(d2, 'Tổng quan tài liệu của chúng tôi cho thấy nhiều nghiên cứu tại Việt Nam sử dụng DASS-21 với ngưỡng cắt thấp, cho tỷ lệ rất cao (40\u201386%). Nghiên cứu này của VinUni \u2014 cơ sở đào tạo uy tín \u2014 bổ sung dữ liệu cho Hà Nội sau COVID-19 (thu thập 8\u201312/2023).', bold=True)

bl(d2, 'Dữ liệu bao gồm khối lớp (10\u201312), loại trường (công lập 68%, tư thục 19%, chuyên 13%), giới tính (nam 21,6%, nữ 77,6%, khác 0,8%), sử dụng thuốc và chất kích thích.')

rb(d2, 'Kết quả nghiên cứu định lượng')
bl(d2, 'Tỷ lệ trầm cảm, lo âu, căng thẳng (N = 501, DASS-21):', bold=True)
tbl(d2,
    ['Mức độ', 'Trầm cảm n (%)', 'Lo âu n (%)', 'Căng thẳng n (%)'],
    [['Bình thường', '106 (21,2%)', '69 (13,8%)', '117 (23,4%)'],
     ['Nhẹ', '78 (15,6%)', '33 (6,6%)', '72 (14,4%)'],
     ['Vừa', '151 (30,1%)', '126 (25,1%)', '120 (24,0%)'],
     ['Nặng', '61 (12,2%)', '87 (17,4%)', '126 (25,1%)'],
     ['Rất nặng', '105 (21,0%)', '186 (37,1%)', '66 (13,2%)'],
     ['Tổng có triệu chứng', '395 (78,8%)', '432 (86,2%)', '384 (76,6%)'],
     ['Điểm TB \u00b1 SD', '17,71 \u00b1 10,14', '16,58 \u00b1 8,45', '22,14 \u00b1 9,04']],
    widths=[3.0, 3.0, 3.0, 3.0])
d2.add_paragraph()

bl(d2, 'So sánh tỷ lệ lo âu với các NC tại Việt Nam:', bold=True)
tbl(d2,
    ['Nghiên cứu', 'Công cụ', 'Lo âu', 'So với BQ'],
    [['Bảo Quyên 2025 (bài này)', 'DASS-21', '86,2%', 'CAO NHẤT'],
     ['Thảo Vi 2025 (Huế)', 'DASS-21', '65,8%', '\u00d70,76'],
     ['An Giang 2025', 'DASS-21', '61,2%', '\u00d70,71'],
     ['Ngô Anh Vinh 2024 (DTTS)', 'DASS-21', '54,4%', '\u00d70,63'],
     ['Danh Lâm 2022 (Thanh Hóa)', 'DASS-21', '49,0%', '\u00d70,57'],
     ['Hoa 2024 (GAD-7)', 'GAD-7', '40,6%', '\u00d70,47'],
     ['Vĩnh Lộc 2024 (DASS-Y)', 'DASS-Y', '25,1%', '\u00d70,29'],
     ['V-NAMHS 2022 (DISC-5)', 'DISC-5', '2,3%', '\u00d70,03']],
    widths=[4.0, 2.0, 2.0, 3.0])
d2.add_paragraph()

bh3(d2, 'Khác biệt giới tính : Y văn quốc tế xác nhận nữ > nam về lo âu (McLean, 2011). Nghiên cứu này phù hợp \u2014 giới tính liên quan có ý nghĩa cả 3 chỉ số (p < 0,05). Tuy nhiên, mẫu có 78% nữ \u2014 thiên lệch nghiêm trọng \u2014 nên kết quả giới tính có thể phản ánh đặc điểm mẫu hơn là dân số thực.')

rb(d2, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3(d2, 'Phát hiện về tỷ lệ cực cao và vai trò công cụ đo lường')
bl(d2, '*Tỷ lệ lo âu 86,2% \u2014 CAO NHẤT trong tất cả NC tại Việt Nam.* DASS-21 với ngưỡng cắt rất thấp (lo âu \u22654) bao gồm cả mức "nhẹ" \u2014 37,1% ở mức "rất nặng" là đáng lo ngại thực sự. Tuy nhiên, so với V-NAMHS 2022 (DISC-5: 2,3%), chênh lệch 37 lần gợi ý phương pháp quyết định kết quả.')
bl(d2, '*Thiên lệch giới và tự chọn.* Nữ chiếm 78% mẫu \u2014 thiên lệch giới tính nghiêm trọng. Khảo sát trực tuyến \u2014 HS có vấn đề SKTT có thể tham gia nhiều hơn (thiên lệch tự chọn). Cỡ mẫu 501 \u2014 nhỏ cho Hà Nội (so Hoa 2024: 3.910).')
bl(d2, '*Lo âu "rất nặng" 37,1%.* Đây là tỷ lệ rất cao \u2014 cần xác nhận bằng chẩn đoán lâm sàng. Nếu sử dụng ngưỡng GAD-7 \u226510 (như Hoa 2024), tỷ lệ có thể giảm đáng kể.')

rb(d2, 'Kết luận')
bl(d2, 'Dữ liệu của chúng tôi, cho thấy tỷ lệ lo âu 86,2% \u2014 cao nhất trong tất cả nghiên cứu tại Việt Nam \u2014 ở 501 HS THPT Hà Nội, gợi ý rằng hoặc DASS-21 với ngưỡng thấp không phù hợp cho sàng lọc dân số chung, hoặc SKTT HS THPT Hà Nội thực sự ở mức báo động cần can thiệp khẩn cấp. Sự chênh lệch 37 lần giữa DASS-21 (86,2%) và DISC-5 (2,3%) nhấn mạnh nhu cầu nghiên cứu so sánh nhiều công cụ trên cùng mẫu \u2014 một khoảng trống phương pháp luận then chốt.', bold=True)

rh2(d2, 'Phản biện')
bl(d2, 'VinUni \u2014 cơ sở đào tạo uy tín, có IRB từ Vinmec, đánh giá cả 3 chỉ số cùng lúc. Tuy nhiên, lo âu 86,2% CỰC CAO \u2014 DASS-21 ngưỡng \u22654 bao gồm cả "nhẹ". Nữ 78% mẫu \u2014 thiên lệch giới nghiêm trọng, không đại diện HS. Online survey \u2014 thiên lệch tự chọn. n = 501 nhỏ (so Hoa 2024: 3.910). Tạp chí YHCĐ không lập chỉ mục quốc tế. Chỉ thống kê đơn biến + hồi quy \u2014 chưa rõ loại hồi quy.')

rh2(d2, 'Hướng nghiên cứu tiếp theo')
bl(d2, 'Tái lặp với cỡ mẫu lớn hơn, tỷ lệ giới cân bằng. So sánh DASS-21 với GAD-7 hoặc DISC-5 trên cùng mẫu \u2014 xác định ngưỡng cắt phù hợp. Can thiệp tại trường THPT Hà Nội. Xuất bản trên tạp chí có chỉ mục quốc tế.')

d2.add_paragraph()
p = d2.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50 Trung bình-Thấp. VinUni uy tín nhưng tỷ lệ 86,2% cao bất thường, thiên lệch giới 78% nữ, online survey, n = 501 nhỏ, tạp chí không chỉ mục QT.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d2.save('VN16_BaoQuyen_2025.docx')
sys.stderr.write('VN16 OK\n')
