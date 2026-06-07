# -*- coding: utf-8 -*-
"""Check plagiarism Q1 draft vs LA chinh.
Compare Vietnamese mirror paragraphs in Q1 vs LA Vietnamese.
Find 7+ consecutive word matches (Turnitin/Crossref threshold)."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Extract Q1 Vietnamese paragraphs (mirror of English)
d = Document(os.path.join(ROOT, 'bai-bao-Q1', 'Draft_Q1_SongNgu_v4_01062026.docx'))
q1_vn_paras = []
for p in d.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    if any(c in txt for c in 'ăâđêôơưĂÂĐÊÔƠƯáàảãạếềểễệíìỉĩịóòỏõọốồổỗộúùủũụýỳỷỹỵ'):
        q1_vn_paras.append(txt)
print(f'Q1 VN paragraphs: {len(q1_vn_paras)}')
print(f'Q1 VN total words: {sum(len(p.split()) for p in q1_vn_paras)}')

# Extract LA chinh Vietnamese text
d_la = Document(os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx'))
la_paras = []
for p in d_la.paragraphs:
    t = p.text.strip()
    if t and len(t.split()) > 8:  # skip very short
        la_paras.append(t)
la_full = ' '.join(la_paras).lower()
print(f'LA paragraphs: {len(la_paras)}')


def normalize(text):
    """Lowercase + strip punctuation."""
    text = text.lower()
    text = re.sub(r'[^\w\sàảãạáăắằẳẵặâấầẩẫậéẻẽẹèêếềểễệíìỉĩịóòỏõọôốồổỗộơớờởỡợúùủũụưứừửữựýỳỷỹỵđ]', ' ', text)
    text = re.sub(r'\s+', ' ', text).strip()
    return text


def find_ngram_matches(target_text, corpus_text, n=7):
    """Find n-gram (n consecutive words) matches between target and corpus."""
    target_norm = normalize(target_text)
    corpus_norm = normalize(corpus_text)
    target_words = target_norm.split()
    matches = []
    seen = set()
    for i in range(len(target_words) - n + 1):
        ngram = ' '.join(target_words[i:i+n])
        if ngram in seen: continue
        seen.add(ngram)
        if ngram in corpus_norm and len(ngram) > 30:  # avoid trivial
            matches.append(ngram)
    return matches


print()
print('=== PLAGIARISM CHECK: Q1 VN paragraphs vs LA chinh ===')
print('Threshold: 7+ consecutive words match (Turnitin/Crossref standard)')
print()

total_para = 0
flagged = 0
issues = []

for i, para in enumerate(q1_vn_paras):
    if len(para.split()) < 15: continue  # skip short paragraphs (likely titles)
    total_para += 1
    matches = find_ngram_matches(para, la_full, n=7)
    if matches:
        flagged += 1
        issues.append({
            'para_idx': i,
            'para': para[:200],
            'matches': matches,
        })

print(f'Total Q1 VN paragraphs checked: {total_para}')
print(f'Flagged with 7-gram matches in LA: {flagged}')
if total_para > 0:
    risk_ratio = flagged / total_para
    risk = 'HIGH' if risk_ratio > 0.3 else ('MEDIUM' if flagged > 5 else 'LOW')
else:
    risk = 'N/A'
print(f'Risk level: {risk}')
print()
print(f'=== TOP {min(15, len(issues))} FLAGGED PARAGRAPHS ===')
for j, item in enumerate(issues[:15]):
    para_idx = item['para_idx']
    para_preview = item['para']
    n_matches = len(item['matches'])
    print(f'\n[{j+1}] Para #{para_idx} ({n_matches} matches):')
    print(f'  Q1: {para_preview}...')
    for m in item['matches'][:3]:
        print(f'  -> MATCH: ...{m}...')

# Save full report
report_path = os.path.join(ROOT, '06_Scripts', '_plagiarism_report_Q1_vs_LA.txt')
with open(report_path, 'w', encoding='utf-8') as f:
    f.write(f'PLAGIARISM CHECK: Q1 Vietnamese paragraphs vs LA chinh\n')
    f.write(f'Total checked: {total_para}, Flagged: {flagged}\n\n')
    for j, item in enumerate(issues):
        idx = item['para_idx']
        n = len(item['matches'])
        f.write('\n[' + str(j+1) + '] Q1 Para #' + str(idx) + ':\n')
        f.write('Text: ' + item['para'] + '\n')
        f.write('Matches (' + str(n) + '):\n')
        for m in item['matches']:
            f.write('  - ' + m + '\n')

print(f'\\nFull report saved: {report_path}')
