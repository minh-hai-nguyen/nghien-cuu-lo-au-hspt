# -*- coding: utf-8 -*-
"""
v5 -> v6: sửa fabrication trích dẫn QUỐC TẾ trong 2 bài VJES.
Dữ liệu thư mục đã verify qua web (PubMed/PMC/publisher) 16/05/2026.
- Thay 5 mục TLTK QT Bài 1 (Chen, Xu, Jagiello, Fassi, Lee) + sửa 6 đoạn thân bài.
- Thay 12 mục TLTK QT Bài 2 + XÓA 3 mục (Brown-Carter bịa, Qiaochu/Samele orphan)
  + sửa 9 đoạn thân bài + thêm trích dẫn Li ở khuyến nghị.
v4, v5 GIỮ NGUYÊN để đối chứng.
"""
import os
from docx import Document
from datetime import datetime

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn"
B1_IN = os.path.join(ROOT, "Bai1_YTNC_HSTHCS_v5_15052026.docx")
B2_IN = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v5_15052026.docx")
B1_OUT = os.path.join(ROOT, "Bai1_YTNC_HSTHCS_v6_16052026.docx")
B2_OUT = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v6_16052026.docx")


def replace_para_keep_format(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def replace_by_unique(doc, sub, new_text):
    for p in doc.paragraphs:
        if sub in p.text:
            replace_para_keep_format(p, new_text)
            return True
    return False


def replace_substr(doc, sub, new_sub):
    """Thay 1 cụm con trong đoạn (giữ phần còn lại)."""
    for p in doc.paragraphs:
        if sub in p.text:
            replace_para_keep_format(p, p.text.replace(sub, new_sub))
            return True
    return False


def delete_by_unique(doc, sub):
    for p in doc.paragraphs:
        if sub in p.text:
            p._element.getparent().remove(p._element)
            return True
    return False


# ============================================================
# BÀI 1
# ============================================================
B1_BODY = [
    # [8] — Chen: lo âu 13,9% (không phải 20%); bỏ claim "khối đầu THCS rủi ro cao nhất"
    ("Chen và cộng sự (2023) tại Trung Quốc cho thấy tỷ lệ có triệu chứng lo âu lên tới hai phần mười",
     "Tại khu vực Á – đặc biệt là Đông Á – áp lực học tập, văn hóa thi cử nặng nề và kỳ vọng gia đình được xác định là những đặc trưng làm gia tăng nguy cơ lo âu ở học sinh trung học cơ sở (HSTHCS). Nghiên cứu cắt ngang trên 63.205 học sinh trung học của Chen và cộng sự (2023) tại Trung Quốc ghi nhận tỷ lệ học sinh có triệu chứng lo âu là 13,9%. Khảo sát quy mô lớn của Xu và cộng sự (2021) trên 373.216 học sinh trung học Trung Quốc trong giai đoạn đỉnh COVID-19 ghi nhận tỷ lệ lo âu (thang Rối loạn Lo âu Tổng quát 7 mục, GAD-7 ≥ 10) ở mức 9,89%, trong đó khu vực nông thôn (11,33%) cao hơn khu vực đô thị (8,77%) và bậc trung học cơ sở (13,89%) cao hơn bậc trung học phổ thông (12,93%)."),

    # [19] — Pascoe: bỏ "khung sáu chiều" (bịa); Pascoe là tổng quan TÁC ĐỘNG của stress
    ("Pascoe và cộng sự (2020) đề xuất khung sáu chiều mô tả áp lực học tập",
     "Áp lực học tập từ lâu được xếp vào trục yếu tố nguy cơ trung tâm đối với rối loạn lo âu ở học sinh phổ thông. Tổng quan của Pascoe và cộng sự (2020) về tác động của căng thẳng học đường cho thấy áp lực học tập có liên hệ với suy giảm kết quả học tập, rối loạn giấc ngủ và gia tăng triệu chứng lo âu – trầm cảm ở học sinh trung học và sinh viên. Áp lực học tập không phải là một biến đơn nhất mà là tổ hợp của nhiều thành phần — khối lượng bài tập, kỳ vọng về kết quả thi, lo âu thi cử, cạnh tranh điểm số và kỳ vọng từ giáo viên cũng như phụ huynh; do đó hiệu ứng tích lũy của áp lực học tập trên sức khỏe tâm thần có thể bị đánh giá thấp khi chỉ đo bằng một câu hỏi tự đánh giá đơn lẻ."),

    # [20] — Jagiello: tạp chí đúng; bỏ "(GBD ASEAN, 2025)" (cite ma, không có TLTK)
    ("Jagiello và cộng sự (2025) công bố trên Child Psychiatry and Human Development",
     "Tổng quan hệ thống của Jagiello và cộng sự (2025) công bố trên tạp chí Child Psychiatry & Human Development đã rà soát 31 nghiên cứu can thiệp áp lực học tập ở học sinh trung học tại 13 quốc gia và xác nhận rằng áp lực học tập là một mục tiêu can thiệp khả thi khi triển khai trong môi trường trường học, trong đó các chương trình dựa trên liệu pháp nhận thức – hành vi cho bằng chứng hiệu quả mạnh nhất. Điều này cho thấy áp lực học tập không chỉ là một yếu tố nguy cơ có thể đo lường mà còn là một điểm tác động thực tế cho các chương trình dự phòng tại trường học."),

    # [24] — Fassi: bài thật là nghiên cứu cắt ngang (không phải meta-analysis); g=0,46 đã verify
    ("Nghiên cứu cắt ngang đại diện quốc gia Anh của Fassi và cộng sự (2025)",
     "Nghiên cứu đại diện quốc gia của Vương quốc Anh do Fassi và cộng sự (2025) công bố trên tạp chí Nature Human Behaviour, thực hiện trên 3.340 vị thành niên 11–19 tuổi có đánh giá chẩn đoán lâm sàng, cho thấy nhóm có rối loạn sức khỏe tâm thần được chẩn đoán dành nhiều thời gian sử dụng mạng xã hội hơn nhóm không có rối loạn, với chênh lệch ở mức trung bình (g = 0,46); riêng nhóm có rối loạn dạng hướng nội — bao gồm lo âu — đạt mức chênh cao hơn (g = 0,62). Điểm mạnh đặc biệt của thiết kế này so với các nghiên cứu chỉ dựa trên sàng lọc tự báo cáo là việc tách rõ nhóm có rối loạn được chẩn đoán lâm sàng so với nhóm chứng, qua đó giảm thiên lệch tự lựa chọn."),

    # [33] — Lee: r=0,23 là tương quan TỔNG (mọi triệu chứng SKTT), không riêng lo âu
    ("Phân tích tổng hợp của Lee và cộng sự (2025) trên Trauma, Violence, & Abuse",
     "Bắt nạt bằng lời nói và bắt nạt mạng (cyberbullying) hợp thành nhóm hai dạng bắt nạt phi thể chất ngày càng được giới nghiên cứu chú ý. Trong khi bắt nạt thể chất giảm tỷ lệ tại nhiều khu vực, bắt nạt bằng lời và bắt nạt mạng lại có xu hướng ngược lại do thiếu rào cản kỹ thuật và đặc tính ẩn danh, kéo dài liên tục và có đông người chứng kiến trong môi trường trực tuyến. Phân tích tổng hợp các nghiên cứu dọc của Lee và cộng sự (2025) công bố trên tạp chí Trauma, Violence, & Abuse, tổng hợp 27 nghiên cứu, báo cáo hệ số tương quan tổng hợp giữa nạn nhân bắt nạt mạng và các triệu chứng sức khỏe tâm thần — bao gồm lo âu và trầm cảm — ở mức r = 0,23 (khoảng tin cậy 95% từ 0,19 đến 0,27); kích thước hiệu ứng nhỏ về mặt thống kê nhưng có ý nghĩa lâm sàng đáng kể do tỷ lệ phơi nhiễm cao."),

    # [38] — Sowislo & Orth: nhánh phân tích lo âu dựa trên 18 mẫu (77 mẫu là nhánh trầm cảm)
    ("Sowislo và Orth (2013) trong phân tích tổng hợp các nghiên cứu dọc trên 77 mẫu",
     "Lòng tự trọng thấp được bổ sung như yếu tố thứ năm dựa trên hai cân nhắc: (i) lòng tự trọng vừa là yếu tố nguy cơ độc lập vừa là yếu tố trung gian kết nối các yếu tố tình huống (áp lực học tập, bắt nạt) với triệu chứng lo âu; (ii) các bằng chứng phân tích tổng hợp về lòng tự trọng đã đạt mức tin cậy cao. Sowislo và Orth (2013) trong phân tích tổng hợp các nghiên cứu dọc báo cáo rằng lòng tự trọng thấp dự báo ngược triệu chứng lo âu trong tương lai với hệ số β = −0,10 — nhánh phân tích riêng cho lo âu dựa trên 18 mẫu dọc — ngay cả sau khi kiểm soát đầy đủ mức cơ sở của lo âu. Hệ số tuy nhỏ nhưng nhất quán xuyên các mẫu và thỏa các tiêu chí Bradford Hill về tính nhân – quả khi đặt cùng các bằng chứng khác."),
]

B1_TLTK = [
    ("Chen, Y., Wang, H., & Liu, M. (2023)",
     "Chen, Z., Ren, S., He, R., Liang, Y., Tan, Y., Liu, Y., Wang, F., Shao, X., Chen, S., Liao, Y., He, Y., Li, J.-G., Chen, X., & Tang, J. (2023). Prevalence and associated factors of depressive and anxiety symptoms among Chinese secondary school students. BMC Psychiatry, 23, Article 580. https://doi.org/10.1186/s12888-023-05068-1"),

    ("Xu, D.-D., Rao, W.-W., Cao, X.-L.",
     "Xu, Q., Mao, Z., Wei, D., Liu, P., Fan, K., Wang, J., Wang, X., Hu, F., Hou, Y., & Mao, S. (2021). Prevalence and risk factors for anxiety symptoms during the outbreak of COVID-19: A large survey among 373216 junior and senior high school students in China. Journal of Affective Disorders, 288, 17–22. https://doi.org/10.1016/j.jad.2021.03.080"),

    ("Jagiello, T., Le Couteur, A., & Howard, K. (2025)",
     "Jagiello, T., Belcher, J., Neelakandan, A., Boyd, K., & Wuthrich, V. M. (2025). Academic stress interventions in high schools: A systematic literature review. Child Psychiatry & Human Development, 56(6), 1836–1869. https://doi.org/10.1007/s10578-024-01667-5"),

    ("Fassi, L., Thomas, J., Sala, A., Orben, A., & Etchells, P. J. (2025)",
     "Fassi, L., Ferguson, A. M., Przybylski, A. K., Ford, T. J., & Orben, A. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9(6), 1283–1299. https://doi.org/10.1038/s41562-025-02134-4"),

    ("Lee, C. S., Tang, D. T., Tsai, C. T., Wong, C. S., & Tsai, C. C. (2025)",
     "Lee, J., Choo, H., Zhang, Y., Cheung, H. S., Zhang, Q., & Ang, R. P. (2025). Cyberbullying victimization and mental health symptoms among children and adolescents: A meta-analysis of longitudinal studies. Trauma, Violence, & Abuse. Xuất bản trực tuyến trước. https://doi.org/10.1177/15248380241313051"),
]


# ============================================================
# BÀI 2
# ============================================================
B2_BODY = [
    # [9] §1 — Liu (người lớn, bỏ SMD cụ thể), Xian (SUCRA), Cai (resilience không phải lo âu)
    ("Phân tích mạng lưới của Liu và cộng sự (2025) trên 52 thử nghiệm",
     "Phân tích tổng hợp lớn của Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành niên xác nhận chiến lược đối phó tập trung vào vấn đề – một thành tố cốt lõi của CBT – có liên hệ ngược chiều và đáng kể với triệu chứng lo âu. Tổng quan hệ thống và phân tích mạng lưới của Liu và cộng sự (2025) trên 52 thử nghiệm lâm sàng ngẫu nhiên ở bệnh nhân rối loạn lo âu tổng quát — chủ yếu là người trưởng thành — so sánh các hình thức cung cấp CBT cho thấy CBT cá nhân vượt trội rõ rệt so với nhóm chứng đợi, CBT nhóm cũng vượt trội so với nhóm chứng đợi, trong khi CBT từ xa không cho thấy ưu thế so với nhóm chứng; kết quả này hữu ích như một tham chiếu về thứ bậc hiệu quả giữa các hình thức cung cấp. Phân tích mạng lưới riêng cho rối loạn lo âu xã hội ở trẻ em – vị thành niên của Xian và cộng sự (2024) báo cáo kết quả cùng hướng, trong đó các hình thức CBT — đặc biệt là CBT qua internet — được xếp hạng hiệu quả cao nhất. Trên phương diện học đường, phân tích tổng hợp của Cai và cộng sự (2025) ghi nhận các chương trình tại trường có hiệu ứng tổng hợp khiêm tốn trên năng lực phục hồi tâm lý của học sinh — một cấu trúc bảo vệ có liên quan đến việc giảm nguy cơ lo âu."),

    # [10] §1 — Walder: phạm vi đúng là rối loạn lo âu xã hội
    ("Walder và cộng sự (2025) trên 21 thử nghiệm báo cáo hiệu ứng tổng hợp g = 0,508",
     "Trong bối cảnh chuyển đổi số mạnh mẽ, can thiệp CBT qua nền tảng số (Digital Mental Health Intervention – DMHI) ngày càng được chú ý. Phân tích tổng hợp của Walder và cộng sự (2025) trên 21 thử nghiệm lâm sàng ngẫu nhiên về can thiệp số cho rối loạn lo âu xã hội ở trẻ em, vị thành niên và thanh niên trẻ báo cáo hiệu ứng tổng hợp g = 0,508, trong đó nhóm can thiệp số có hướng dẫn đạt g = 0,825. Tại Việt Nam, các nghiên cứu can thiệp dựa trên CBT hiện có chủ yếu mang tính thử nghiệm pilot quy mô nhỏ hoặc nhắm vào nhóm tuổi khác; các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên HSTHCS Việt Nam vẫn chưa được triển khai."),

    # [22] §3.1 — Liu: bỏ "47 thử nghiệm/g=0,82/0,68/0,55" (bịa); nói đúng là mẫu người lớn
    ("Phân tích mạng lưới của Liu và cộng sự (2025) trên 47 thử nghiệm",
     "Bằng chứng về hiệu quả so sánh giữa các hình thức cung cấp CBT chủ yếu đến từ nghiên cứu trên người trưởng thành. Tổng quan hệ thống và phân tích mạng lưới của Liu và cộng sự (2025) trên 52 thử nghiệm lâm sàng ngẫu nhiên ở bệnh nhân rối loạn lo âu tổng quát cho thấy CBT cá nhân đạt hiệu quả cao nhất, vượt trội so với CBT từ xa, điều trị thông thường và nhóm chứng đợi; CBT nhóm cũng vượt trội so với nhóm chứng đợi, trong khi CBT từ xa không cho thấy ưu thế rõ so với điều trị thông thường hay nhóm chứng đợi. Mặc dù mẫu nghiên cứu là người trưởng thành, thứ bậc hiệu quả này là một tham chiếu hữu ích khi lựa chọn hình thức triển khai CBT cho học sinh, đồng thời cho thấy hiệu quả của một giao thức phụ thuộc đáng kể vào việc nó được thực hiện đầy đủ và có hỗ trợ chuyên môn hay không."),

    # [23] §3.1 — Xian: bỏ "g=0,7-0,9" (bịa), bài thật báo cáo SUCRA; Praptomojati là tổng quan hệ thống
    ("Xian và cộng sự (2024) xác nhận CBT là liệu pháp tâm lý có hiệu quả cao nhất",
     "Đối với rối loạn lo âu xã hội – một trong những dạng phổ biến ở giai đoạn đầu vị thành niên – phân tích mạng lưới của Xian và cộng sự (2024) trên 30 thử nghiệm lâm sàng ngẫu nhiên ở trẻ em và vị thành niên xác định các hình thức CBT chiếm những vị trí xếp hạng hiệu quả cao nhất, trong đó CBT qua internet đứng đầu (chỉ số SUCRA 71,2%), tiếp đến là CBT nhóm và CBT cá nhân. Praptomojati và cộng sự (2024) trong tổng quan hệ thống các phiên bản CBT thích nghi văn hóa (CA-CBT) cho rối loạn lo âu tại Đông Nam Á nhận định rằng các yếu tố như tôn giáo, gia đình mở rộng và quan niệm cộng đồng về sức khỏe tâm thần cần được tích hợp tường minh vào giao thức để đạt hiệu quả cao nhất ở khu vực này. Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành niên xác nhận rằng đối phó tập trung vào vấn đề và đánh giá lại nhận thức có liên hệ ngược chiều và bền vững với triệu chứng lo âu – trầm cảm, khẳng định nền tảng lý thuyết của CBT."),

    # [29] §3.2 — Cai: ghi rõ effect size là của NĂNG LỰC PHỤC HỒI, không phải lo âu
    ("Phân tích tổng hợp của Cai và cộng sự (2025) trên 38 thử nghiệm",
     "Phân tích tổng hợp của Cai và cộng sự (2025) trên 38 thử nghiệm lâm sàng ngẫu nhiên (21 thử nghiệm trong phân tích định lượng) về các chương trình tăng cường năng lực phục hồi tâm lý tại trường báo cáo hiệu ứng tổng hợp trên năng lực phục hồi đạt SMD = 0,17 (khoảng tin cậy 95% từ 0,06 đến 0,29). Khi phân nhóm, các chương trình dựa trên chánh niệm cho hiệu ứng cao nhất (SMD = 0,57), tiếp đến là các chương trình lấy hoạt động thể thao làm trung tâm (SMD = 0,41). Mặc dù hiệu ứng tổng hợp nhỏ, các tác giả nhấn mạnh giá trị của can thiệp phổ quát quy mô lớn không qua sàng lọc – đặc biệt khi được tích hợp vào hoạt động trải nghiệm hoặc môn Giáo dục công dân; năng lực phục hồi được củng cố là một cấu trúc bảo vệ có liên quan đến việc giảm nguy cơ rối loạn lo âu."),

    # [30] §3.2 — Bradshaw: 12 buổi phụ huynh (không phải 16); theo dõi đến hết lớp 8 (~1 năm)
    ("Thử nghiệm lâm sàng ngẫu nhiên cụm của Bradshaw, Lochman và cộng sự (2025)",
     "Thử nghiệm lâm sàng ngẫu nhiên cụm của Bradshaw và cộng sự (2025) đánh giá chương trình Early Adolescent Coping Power (EACP) trên 709 học sinh lớp 7 tại 40 trường trung học cơ sở của Hoa Kỳ. Chương trình rút gọn gồm 25 buổi với học sinh và 12 buổi với cha mẹ, tích hợp các thành phần CBT cốt lõi. Theo dõi đến cuối lớp 8 — khoảng một năm sau can thiệp — cho thấy nhóm EACP giảm các vấn đề ngoại hiện do giáo viên đánh giá so với nhóm chứng, đồng thời ghi nhận cải thiện có ý nghĩa ở một số kết cục đối với học sinh nữ. Mặc dù đối tượng gốc của EACP là học sinh có hành vi gây hấn chứ không phải lo âu thuần, các thành phần CBT cốt lõi có thể được bản địa hóa để hướng tới lo âu trong bối cảnh Việt Nam."),

    # [31] §3.2 — He: lo âu giảm KHÔNG có ý nghĩa thống kê; XÓA Brown-Carter (bịa)
    ("Urao và cộng sự (2018, 2022) phát triển chương trình Journey of the Brave",
     "Tại Nhật Bản, Urao và cộng sự (2018, 2022) phát triển chương trình Journey of the Brave dựa trên CBT — phiên bản gốc 10 buổi (Urao và cộng sự, 2018) sau đó được mở rộng thành 14 buổi hoạt động ngắn trong lớp học (Urao và cộng sự, 2022) — thiết kế cho học sinh cuối tiểu học và được kiểm định qua các nghiên cứu tựa thực nghiệm có nhóm chứng, ghi nhận hiệu ứng giảm triệu chứng lo âu. Tại Trung Quốc, He và cộng sự (2026) báo cáo một nghiên cứu thử nghiệm chương trình CBT trường học ngắn (Power Up – CBTD) cho học sinh có triệu chứng trầm cảm, công bố trên tạp chí Journal of Affective Disorders; chương trình cải thiện triệu chứng trầm cảm, trong khi triệu chứng lo âu có xu hướng giảm nhưng khác biệt giữa hai nhóm chưa đạt ý nghĩa thống kê — cho thấy hiệu quả lan tỏa sang lo âu của các can thiệp tập trung vào trầm cảm vẫn cần được kiểm chứng thêm."),

    # [32] §3.2 — bỏ claim "Cai: thành phần phụ huynh hiệu ứng cao hơn" (không verify được)
    ("Phân tích của Cai và cộng sự (2025) còn cho thấy các chương trình có thành phần phụ huynh",
     "Tại Việt Nam, các chương trình CBT tại trường còn ở giai đoạn rất sớm. Báo cáo của UNICEF Việt Nam (2022) ghi nhận phần lớn trường được khảo sát có ít nhất một hoạt động hỗ trợ tâm lý học đường, nhưng đa số là tham vấn cá nhân khi học sinh chủ động tìm đến, không phải chương trình phổ quát có cấu trúc. Khoảng trống lớn nhất tại Việt Nam là sự thiếu vắng một chương trình CBT bản địa được kiểm định trên học sinh trung học cơ sở, thiết kế phù hợp với khung chương trình môn Hoạt động trải nghiệm – Hướng nghiệp của Chương trình giáo dục phổ thông 2018. Một mô hình hai tầng có thể tham khảo: tầng phổ quát (4–6 buổi tích hợp môn học chính khóa, tập trung vào kỹ năng nhận diện cảm xúc) và tầng có chọn lọc (8–12 buổi cho học sinh được sàng lọc có triệu chứng cận lâm sàng). Việc huy động sự đồng hành của phụ huynh trong quá trình can thiệp cũng là một yếu tố cần được cân nhắc khi thiết kế chương trình cho bối cảnh văn hóa Việt Nam."),

    # [36] §3.3 — Walder: g không hướng dẫn = 0,27 (không phải 0,30); phạm vi lo âu xã hội
    ("Walder và cộng sự (2025) trên 21 thử nghiệm lâm sàng ngẫu nhiên báo cáo hiệu ứng tổng hợp g = 0,508 cho DMHI",
     "Phân tích tổng hợp của Walder và cộng sự (2025) trên 21 thử nghiệm lâm sàng ngẫu nhiên về can thiệp số cho rối loạn lo âu xã hội ở trẻ em, vị thành niên và thanh niên trẻ báo cáo hiệu ứng tổng hợp g = 0,508 (khoảng tin cậy 95% từ 0,31 đến 0,71). Khi phân nhóm theo mức độ hỗ trợ, các can thiệp số có hướng dẫn đạt hiệu ứng g = 0,825, cao hơn rõ rệt so với các can thiệp không hướng dẫn (g = 0,27) – cho thấy thành phần hỗ trợ chuyên môn, dù tối thiểu, vẫn đóng vai trò quan trọng đối với hiệu quả của can thiệp số."),

    # [37] §3.3 — Matsumoto: tạp chí JMIR Pediatrics and Parenting; 10 tuần (không phải 8); dưới ngưỡng
    ("Matsumoto và cộng sự (2024) phát triển ứng dụng iCBT",
     "Matsumoto và cộng sự (2024) đánh giá một chương trình CBT qua internet không hướng dẫn trên điện thoại cho vị thành niên và thanh niên trẻ Nhật Bản có rối loạn lo âu xã hội dưới ngưỡng, công bố trên tạp chí JMIR Pediatrics and Parenting. Thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu 77 người tham gia 15–25 tuổi báo cáo hiệu ứng g = 0,66 đối với điểm lo âu xã hội sau 10 tuần can thiệp – con số đáng chú ý vì can thiệp gần như không có sự tham gia trực tiếp của chuyên gia. Yếu tố thiết kế tương tác được xác định là chìa khóa duy trì sự tham gia của người dùng."),
]

# [46] — thêm trích dẫn Maya/ClearlyMe (để Li không thành orphan)
B2_SUBSTR = [
    ("theo mô hình của Maya hoặc ClearlyMe",
     "theo mô hình của các ứng dụng đã được phát triển trên thế giới như Maya (Bress và cộng sự, 2024) hoặc ClearlyMe (Li và cộng sự, 2022)"),
]

B2_TLTK = [
    ("Liu, P., Han, H., Chen, J., & Wang, Y. (2025)",
     "Liu, S., Xiao, H., Duan, Y., Shi, L., Wang, P., Cao, L., Li, H., Huang, X., & Qiu, C. (2025). CBT treatment delivery formats for generalized anxiety disorder: A systematic review and network meta-analysis of randomized controlled trials. Translational Psychiatry, 15, Article 197. https://doi.org/10.1038/s41398-025-03414-3"),

    ("Walder, N., Berger, T., & Krieger, T. (2025)",
     "Walder, N., Frey, A., Berger, T., & Schmidt, S. J. (2025). Digital mental health interventions for the prevention and treatment of social anxiety disorder in children, adolescents, and young adults: A systematic review and meta-analysis of randomized controlled trials. Journal of Medical Internet Research, 27, Article e67067. https://doi.org/10.2196/67067"),

    ("Cai, Y., Wang, H., Liu, M., & Chen, X. (2025)",
     "Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, Article 1594658. https://doi.org/10.3389/fpsyt.2025.1594658"),

    ("Xian, J., Tian, M., Liu, X., & Sun, Y. (2024)",
     "Xian, J., Zhang, Y., & Jiang, B. (2024). Psychological interventions for social anxiety disorder in children and adolescents: A systematic review and network meta-analysis. Journal of Affective Disorders, 365, 614–627. https://doi.org/10.1016/j.jad.2024.08.097"),

    ("Bradshaw, C. P., Lochman, J. E., Powell, N. P., Boxmeyer, C. L.",
     "Bradshaw, C. P., McDaniel, H. L., Pas, E. T., Debnam, K. J., Bottiani, J. H., Powell, N., Ialongo, N. S., Morgan-Lopez, A. A., & Lochman, J. E. (2025). Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology, 110, Article 101437. https://doi.org/10.1016/j.jsp.2025.101437"),

    ("He, X., Liu, Y., Zhou, J., & Wang, Q. (2025)",
     "He, Q., Li, J., Wang, J., & Qu, Z. (2026). Preventing depression in Chinese children and adolescents: A pilot study of a brief school-based cognitive behavioral group program. Journal of Affective Disorders, 394, Article 120559. https://doi.org/10.1016/j.jad.2025.120559"),

    ("Matsumoto, K., Hamatani, S., Nagai, K., Sutoh, C., Nakagawa, A., & Shimizu, E. (2024)",
     "Matsumoto, K., Hamatani, S., Shiga, K., Iiboshi, K., Kasai, M., Kimura, Y., Yokota, S., Watanabe, K., Kubo, Y., & Nakamura, M. (2024). Effectiveness of unguided internet-based cognitive behavioral therapy for subthreshold social anxiety disorder in adolescents and young adults: Multicenter randomized controlled trial. JMIR Pediatrics and Parenting, 7, Article e55786. https://doi.org/10.2196/55786"),

    ("Bress, J. N., Stewart, C., Fenster, R. J., Ehrlich, A. L.",
     "Bress, J. N., Falk, A., Schier, M. M., Jaywant, A., Moroney, E., Dargis, M., Bennett, S. M., Scult, M. A., Volpp, K. G., Asch, D. A., Balachandran, M., Perlis, R. H., Lee, F. S., & Gunning, F. M. (2024). Efficacy of a mobile app-based intervention for young adults with anxiety disorders: A randomized clinical trial. JAMA Network Open, 7(8), Article e2428372. https://doi.org/10.1001/jamanetworkopen.2024.28372"),

    ("Praptomojati, A., Suryani, A. O., Salim, R. M. A., & Pratiwi, T. F. (2024)",
     "Praptomojati, A., Icanervilia, A. V., Nauta, M. H., & Bouman, T. K. (2024). A systematic review of culturally adapted cognitive behavioral therapy (CA-CBT) for anxiety disorders in Southeast Asia. Asian Journal of Psychiatry, 92, Article 103896. https://doi.org/10.1016/j.ajp.2023.103896"),

    ("Li, S. H., Cheng, J. W., Yang, M., Newby, J. M., & Christensen, H. (2024)",
     "Li, S. H., Achilles, M. R., Werner-Seidler, A., Beames, J. R., Subotic-Kerry, M., O'Dea, B., & Christensen, H. (2022). A cognitive behavioural therapy smartphone app for adolescent depression and anxiety: Co-design of ClearlyMe. The Cognitive Behaviour Therapist, 15, Article e21. https://doi.org/10.1017/S1754470X22000095"),

    ("Urao, Y., Yoshida, M., Koshiba, T., Sato, Y., Ishikawa, S., & Shimizu, E. (2018)",
     "Urao, Y., Yoshida, M., Koshiba, T., Sato, Y., Ishikawa, S., & Shimizu, E. (2018). Effectiveness of a cognitive behavioural therapy-based anxiety prevention programme at an elementary school in Japan: A quasi-experimental study. Child and Adolescent Psychiatry and Mental Health, 12, Article 33. https://doi.org/10.1186/s13034-018-0240-5"),

    ("Urao, Y., Yoshinaga, N., Asano, K., Ishikawa, R., Tano, A., Sato, Y., Shimizu, E. (2022)",
     "Urao, Y., Yoshida, M., Sato, Y., & Shimizu, E. (2022). School-based cognitive behavioural intervention programme for addressing anxiety in 10- to 11-year-olds using short classroom activities in Japan: A quasi-experimental study. BMC Psychiatry, 22, Article 658. https://doi.org/10.1186/s12888-022-04326-y"),
]

# Xóa 3 mục TLTK: Brown-Carter (bịa hoàn toàn), Qiaochu + Samele (orphan, không cite trong bài)
B2_TLTK_DELETE = [
    "Brown-Carter, J. L., Patalay, P., & Stallard, P. (2025)",
    "Qiaochu, Z., Wang, H., & Zhang, M. (2025)",
    "Samele, C., Stallard, P., Verkuyl, M., & Tudor, A. (2025)",
]


def process(in_path, out_path, body, tltk, label, substr=None, tltk_del=None):
    print(f"=== {label} ===")
    doc = Document(in_path)
    nb = nt = ns = nd = 0
    for sub, new in body:
        if replace_by_unique(doc, sub, new):
            nb += 1
        else:
            print(f"  !! BODY NOT FOUND: {sub[:55]}")
    for sub, new in tltk:
        if replace_by_unique(doc, sub, new):
            nt += 1
        else:
            print(f"  !! TLTK NOT FOUND: {sub[:55]}")
    if substr:
        for sub, new in substr:
            if replace_substr(doc, sub, new):
                ns += 1
            else:
                print(f"  !! SUBSTR NOT FOUND: {sub[:55]}")
    if tltk_del:
        for sub in tltk_del:
            if delete_by_unique(doc, sub):
                nd += 1
            else:
                print(f"  !! DELETE NOT FOUND: {sub[:55]}")
    print(f"  body {nb}/{len(body)}, TLTK {nt}/{len(tltk)}, "
          f"substr {ns}/{len(substr) if substr else 0}, deleted {nd}/{len(tltk_del) if tltk_del else 0}")

    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.subject = ""
    cp.category = ""
    cp.title = ""
    cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 16, 10, 0, 0)
    doc.save(out_path)
    print(f"  Saved: {out_path}\n")


if __name__ == "__main__":
    process(B1_IN, B1_OUT, B1_BODY, B1_TLTK, "BAI 1")
    process(B2_IN, B2_OUT, B2_BODY, B2_TLTK, "BAI 2",
            substr=B2_SUBSTR, tltk_del=B2_TLTK_DELETE)
    print("[DONE]")
