# -*- coding: utf-8 -*-
"""B: Them phan ERRATA vao cuoi file ThamKhao_Titles_v2.
Ghi nhan tat ca 6 loi da phat hien + correction.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
INPUT = os.path.join(ROOT, 'bai-bao-Q1', 'ThamKhao_Titles_Q1Q3_AsiaChauPhi_TiengViet_v2_01062026.docx')

d = Document(INPUT)

# Add page break + ERRATA section at end
d.add_page_break()


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def P(text, italic=False, indent=True):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic


H1('ERRATA — CÁC SỬA SAU KHI AUDIT (01/06/2026 — sau khi file v2 đã tạo)')

P('Sau khi tạo file v2, em đã kiểm tra lại từng paper bằng search trực tiếp '
  'trên PubMed, Crossref, và nature.com để verify tác giả, năm, và tạp chí. '
  'Quá trình audit phát hiện 6 lỗi cần sửa. Phần ERRATA này ghi nhận các '
  'lỗi đó để minh bạch và để người đọc tham khảo đúng thông tin.', italic=False)


H2('Lỗi 1 — Paper #1 (Phần 2.1, mục #1)')
P('Em ghi: "Wang và cộng sự — Trung Quốc"', indent=False)
P('Đúng: Juan J, Li J, Wang X et al. (Wang chỉ là tác giả thứ 3, không phải '
  'đầu). Title gốc còn có "...in Chinese adolescent GIRLS" — mẫu chỉ là '
  'nữ giới, em đã ghi sót.', italic=True, indent=False)
P('Citation đúng: Juan, J., Li, J., Wang, X. et al. Risk factors of '
  'depressive and anxiety symptoms in Chinese adolescent girls: a '
  'cross-sectional study. Sci Rep 15, 30891 (2025).', italic=True, indent=False)


H2('Lỗi 2 — Paper #6 (Phần 2.1, mục #6)')
P('Em ghi: "Tác giả Trung Quốc" + "2025"', indent=False)
P('Đúng: Aijun Zhu, Di Xue, Huaijie Yang, Yanfang Ren (China Three Gorges '
  'University, Yichang). Năm xuất bản thực sự là 16/01/2026, không phải '
  '2025. n=3,673 thanh thiếu niên.', italic=True, indent=False)


H2('Lỗi 3 — Paper #7 (Phần 2.1, mục #7) — NGHIÊM TRỌNG')
P('Em ghi: dùng làm ví dụ Q1 cho "thanh thiếu niên"', indent=False)
P('Đúng: Mẫu nghiên cứu là 1,097 SINH VIÊN ĐẠI HỌC từ Hồ Nam (Hunan), '
  'Trung Quốc — KHÔNG phải adolescent (thanh thiếu niên). Paper này thực '
  'ra không nên có trong danh sách paper về adolescent.', italic=True, indent=False)
P('Khuyến nghị: Xem Paper #7 chỉ như ví dụ mẫu cấu trúc title '
  '"moderated mediating model", không dùng nội dung tham khảo cho bài Q1 '
  'của nhóm.', italic=True, indent=False)


H2('Lỗi 4 — Paper #9 (Phần 3.1, mục #9)')
P('Em ghi: "Tác giả Việt Nam (Hà Nội)" + "PMC\\n2023"', indent=False)
P('Đúng: Tran Minh Dien, Pham Thi Lan Chi, Pham Quang Duy, Le Ha Anh, '
  'Nguyen Thi Kim Ngan, Vu Thi Hoang Lan. n=5,325 HS THCS Hà Nội (10-17 '
  'tuổi). Journal cụ thể là BMC Public Health (em ghi chung "PMC").',
  italic=True, indent=False)
P('Citation đúng: Tran MD, Pham TLC, Pham QD, et al. Prevalence of '
  'internet addiction and anxiety, and factors associated with the high '
  'level of anxiety among adolescents in Hanoi, Vietnam during the COVID-19 '
  'pandemic. BMC Public Health. 2023 Dec 6;23(1):2400.', italic=True, indent=False)


H2('Lỗi 5 — Paper #13 (Phần 2.2, mục #13)')
P('Em ghi: "Tác giả Ethiopia (n=1.407)" + "2023"', indent=False)
P('Đúng: Năm xuất bản thực sự là 10/04/2024, không phải 2023. Em đã nhầm '
  'với năm submit (2/2023). Tác giả đầu: Zenebe Abebe Gebreegziabher.',
  italic=True, indent=False)


H2('Lỗi 6 — Paper #17 (Phần 2.2, mục #17)')
P('Em ghi: "Tác giả Trung Quốc (lab Rwanda)"', indent=False)
P('Đúng: Lisa Cynthia Niwenahisemo, Su Hong, Li Kuang (từ Department of '
  'Psychiatry, The First Affiliated Hospital of Chongqing Medical University). '
  'Em ghi "lab Rwanda" không chính xác — họ là Chinese researchers nhưng '
  'sample là 1,813 HS trung học Kigali, Rwanda (12-17 tuổi).',
  italic=True, indent=False)


H2('Tổng quan các lỗi')

P('Trên tổng 33 paper được liệt kê, em phát hiện 6 paper có vấn đề về '
  'tác giả, năm xuất bản, hoặc đối tượng nghiên cứu (chiếm 18,2%). Ngoài '
  'ra, các chỉ số IF và acceptance rate đã được sửa trong bản v2 chính (so '
  'với bản v1 trước). Em đã build file v3 (ThamKhao_Titles_Q1Q3_AsiaChauPhi_'
  'TiengViet_v3_01062026.docx) tích hợp các sửa này vào nội dung chính.',
  italic=True)

P('Em xin lỗi vì các lỗi này và sẽ thận trọng hơn trong các phiên tới — '
  'luôn verify từng entry bằng search trực tiếp trên PubMed/Crossref trước '
  'khi đưa vào tài liệu gửi đi.', italic=True)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(INPUT)
print(f'Saved ERRATA to: {INPUT}')
print(f'Size: {os.path.getsize(INPUT)} bytes')
