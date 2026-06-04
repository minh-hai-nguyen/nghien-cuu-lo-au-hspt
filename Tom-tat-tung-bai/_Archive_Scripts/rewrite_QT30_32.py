# -*- coding: utf-8 -*-
"""QT30 GBD (2t+4h) + QT31 59Countries (4t) + QT32 Ireland (5t)"""
import sys,io
try:sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
except:pass
from docx import Document
from docx.shared import Pt,Cm,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
RED=RGBColor(0xFF,0,0);BLUE=RGBColor(0,0x70,0xC0)
def md():
    d=Document();s=d.styles['Normal'];s.font.name='Times New Roman';s.font.size=Pt(12)
    s.paragraph_format.space_after=Pt(4);s.paragraph_format.line_spacing=1.5
    for sec in d.sections:sec.top_margin=Cm(2.5);sec.bottom_margin=Cm(2.5);sec.left_margin=Cm(3);sec.right_margin=Cm(2)
    return d
def rb(d,t):p=d.add_paragraph();r=p.add_run(t);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=RED
def bl(d,t,bold=False):p=d.add_paragraph();r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=BLUE;r.bold=bold
def rh2(d,t):
    h=d.add_heading(t,level=2)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=RED
def bh3(d,t):
    h=d.add_heading(t,level=3)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=BLUE
def sh(c,co):s=OxmlElement('w:shd');s.set(qn('w:fill'),co);s.set(qn('w:val'),'clear');c._tc.get_or_add_tcPr().append(s)
def sw(c,w):tcW=c._tc.get_or_add_tcPr();we=OxmlElement('w:tcW');we.set(qn('w:w'),str(int(w*567)));we.set(qn('w:type'),'dxa');tcW.append(we)
def tbl(d,headers,rows,widths):
    t=d.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):sw(row.cells[ci],widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
        sh(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci];c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:r.font.name='Times New Roman';r.font.size=Pt(10)

# ===== QT30: GBD TRENDS 10-24y (2 bảng + 4 hình gốc) =====
d=md()
bl(d,'Tóm tắt bài QT-30',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Xu hướng rối loạn trầm cảm và lo âu ở thanh thiếu niên và thanh niên (10\u201324 tuổi) từ 1990 đến 2021: Phân tích từ nghiên cứu Gánh nặng Bệnh tật Toàn cầu \u00bb (Trends in Depressive and Anxiety Disorders among Adolescents and Young Adults Aged 10\u201324 from 1990 to 2021: A GBD Study Analysis), do Zhang Dongjun, Wu Mingyue và cộng sự (2025), Đại học Y khoa Tân Hương, Trung Quốc, xuất bản trên Journal of Affective Disorders, 387:119491 (Q1, IF \u2248 6,6). Dữ liệu GBD 2021 từ 204 nước/vùng lãnh thổ. 14 trang, 2 bảng, 4 hình.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Công trình sử dụng dữ liệu Gánh nặng Bệnh tật Toàn cầu (GBD 2021) phân tích xu hướng tỷ suất mới mắc, tỷ lệ hiện mắc và DALYs cho trầm cảm và lo âu ở nhóm 10\u201324 tuổi. Nói cách khác, đây là phân tích xu hướng 31 năm (1990\u20132021) trên quy mô toàn cầu sử dụng hai phương pháp thống kê nâng cao: hồi quy Joinpoint và phân tích Tuổi-Giai đoạn-Đoàn hệ (APC).',bold=True)
bl(d,'Joinpoint regression xác định các điểm bẻ (inflection points) trong xu hướng dài hạn, ước tính AAPC (Average Annual Percent Change) cho từng giai đoạn. Bất bình đẳng giữa các nước được đánh giá bằng SII (slope index of inequality) và CI (concentration index).')

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 1 gốc). Xu hướng toàn cầu 1990\u20132021:', bold=True)
tbl(d,['Rối loạn','Tỷ suất 1990\n(/100.000)','Tỷ suất 2021\n(/100.000)','AAPC\n(KTC 95%)','Xu hướng'],
    [['Trầm cảm (mới mắc)','3.085,5','3.909,9','0,97 (0,75\u20131,19)','Tăng'],
     ['Lo âu (mới mắc)','708,0','883,1','0,84 (0,48\u20131,21)','Tăng'],
     ['Trầm cảm (DALYs)','\u2014','\u2014','Tăng 2014\u20132021','Đỉnh 2019'],
     ['Lo âu (DALYs)','\u2014','\u2014','Tăng 2014\u20132021','Tăng mạnh sau 2019']],
    widths=[3.0,2.5,2.5,2.5,2.5])
d.add_paragraph()

bl(d,'Bảng 2. Phân tầng theo giới, tuổi, vùng và SDI:', bold=True)
tbl(d,['Phân tầng','Phát hiện chính'],
    [['Giới tính','Nữ luôn cao hơn, nhưng tốc độ tăng trầm cảm ở NAM nhanh hơn'],
     ['Tuổi 10\u201314','Tăng nhanh nhất cho trầm cảm'],
     ['Tuổi 20\u201324','Tăng nhanh nhất cho lo âu'],
     ['SDI cao (nước giàu)','Gánh nặng lớn nhất + tốc độ tăng nhanh nhất'],
     ['Đông Á','Tốc độ tăng CHẬM NHẤT'],
     ['Bắc Mỹ thu nhập cao','AAPC trầm cảm cao nhất: 2,30'],
     ['Mỹ Latinh Andes','AAPC lo âu cao nhất: 1,91'],
     ['COVID-19 (2020\u20132021)','Xu hướng tăng tăng tốc sau 2019']],
    widths=[4.0,8.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : AAPC lo âu 0,84%/năm trong 31 năm nhất quán với JAACAP 2024 (Mỹ: tăng gấp đôi 8 năm), Na Uy 2025 (tăng 13 năm), Ireland 2024 (tăng 7 năm). Đáng chú ý, Đông Á tăng CHẬM NHẤT \u2014 có thể do hệ thống giáo dục tập trung + gia đình gắn kết. Tại VN, Hoàng Trung Học 2025 cho thấy lo âu giảm sau COVID (41,5% \u2192 25,4%) — phù hợp xu hướng phục hồi tại châu Á.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*Tăng tốc sau 2014.* Joinpoint xác định điểm bẻ 2014 \u2014 gánh nặng tăng nhanh từ đó. Trùng với sự phổ biến của smartphone/MXH (phù hợp Nature 2025, Norway 2025).')
bl(d,'*SDI cao gánh nặng lớn nhất.* Nghịch lý: nước giàu có dịch vụ tốt nhưng gánh nặng lo âu cao nhất. Có thể do phát hiện tốt hơn + áp lực xã hội hiện đại.')
bl(d,'*Bất bình đẳng tuyệt đối mở rộng.* SII cho thấy khoảng cách giàu-nghèo về gánh nặng SKTT tăng theo thời gian.')

rb(d,'Kết luận')
bl(d,'Dữ liệu GBD 2021 từ 204 nước, cho thấy gánh nặng trầm cảm và lo âu ở VTN 10\u201324 tuổi tăng liên tục 31 năm (AAPC 0,84\u20130,97%), tăng tốc mạnh sau 2014, với nước SDI cao gánh nặng lớn nhất và Đông Á chậm nhất, gợi ý rằng xu hướng tăng là toàn cầu nhưng tốc độ phụ thuộc bối cảnh kinh tế xã hội \u2014 VN cần theo dõi xu hướng dài hạn để can thiệp kịp thời.',bold=True)

rh2(d,'Phản biện')
bl(d,'J Affect Disord Q1 IF = 6,6. GBD 2021, 204 nước, 31 năm, Joinpoint + APC nâng cao. Tuy nhiên, GBD phụ thuộc dữ liệu báo cáo quốc gia (VN thiếu dữ liệu chất lượng). Tăng có thể phản ánh tăng phát hiện + tăng thực. COVID-19 (2020\u20132021) ảnh hưởng mạnh nhưng chỉ 2 năm dữ liệu.')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'Phân tích riêng ASEAN/VN từ GBD. Giải thích tại sao Đông Á tăng chậm nhất. Đánh giá tác động COVID dài hạn (GBD 2024+).')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. Q1 IF=6,6, GBD 204 nước, 31 năm, Joinpoint+APC.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT30_GBD_Trends.docx');sys.stderr.write('QT30 OK\n')

# ===== QT31: 59 COUNTRIES (4 bảng gốc) =====
d=md()
bl(d,'Tóm tắt bài QT-31',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Tỷ lệ và các yếu tố liên quan của lo âu ở thanh thiếu niên đi học: Phân tích từ 59 quốc gia \u00bb (Prevalence of and Factors Associated with Anxiety among School Going Adolescents: Analysis from 59 Countries), do Islam MA và cộng sự (2025), xuất bản trên Journal of Affective Disorders (Q1, IF \u2248 6,6). Dữ liệu GSHS (Global School-based Student Health Survey) từ 59 quốc gia. 11 trang, 4 bảng.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Công trình sử dụng dữ liệu Khảo sát Sức khỏe Học sinh Toàn cầu Tại Trường học (GSHS \u2014 Global School-based Student Health Survey) do WHO phối hợp CDC Hoa Kỳ triển khai. Nói cách khác, đây là phân tích đa quốc gia lớn nhất về lo âu ở VTN đi học, sử dụng dữ liệu chuẩn hóa quốc tế từ 59 nước.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 2 gốc). Tỷ lệ lo âu theo vùng WHO:', bold=True)
tbl(d,['Vùng WHO','Tỷ lệ lo âu','So sánh'],
    [['Đông Địa Trung Hải','14,6%','CAO NHẤT'],
     ['Châu Phi','11,0%','Thứ 2'],
     ['Châu Mỹ','~8\u201310%','\u2014'],
     ['Châu Âu','~6\u20138%','\u2014'],
     ['Tây Thái Bình Dương','~5\u20138%','\u2014'],
     ['Đông Nam Á','3,6%','THẤP NHẤT']],
    widths=[4.0,3.0,4.0])
d.add_paragraph()

bl(d,'Bảng 2 (từ Table 3 gốc). Yếu tố liên quan đa quốc gia:', bold=True)
tbl(d,['Yếu tố','Hướng tác động','Mức liên quan'],
    [['Bị bắt nạt','Tăng lo âu','Mạnh, nhất quán liên quốc gia'],
     ['Cô đơn','Tăng lo âu','Mạnh'],
     ['Thiếu bạn thân','Tăng lo âu','Trung bình'],
     ['Hỗ trợ cha mẹ','Giảm lo âu (bảo vệ)','Mạnh, nhất quán'],
     ['Giáo dục cha mẹ cao','Giảm lo âu (bảo vệ)','Trung bình']],
    widths=[4.0,3.5,4.5])
d.add_paragraph()

bl(d,'Bảng 3. So sánh tỷ lệ lo âu: GSHS vs NC khác tại cùng vùng:', bold=True)
tbl(d,['Nguồn','Vùng','Tỷ lệ lo âu','Công cụ','Giải thích chênh lệch'],
    [['GSHS (bài này)','ĐNA','3,6%','GSHS (1\u20132 câu)','Ngưỡng cao, ít câu'],
     ['V-NAMHS 2022','VN','2,3% (chẩn đoán)','DISC-5/DSM-5','Chẩn đoán chuẩn'],
     ['Hoa 2024','VN (Hà Nội)','40,6%','GAD-7 (\u22655)','Sàng lọc, ngưỡng thấp'],
     ['Xu 2021','TQ','9,89%','GAD-7 (\u226510)','Ngưỡng trung bình'],
     ['GBD 2025','Toàn cầu','3,6\u20134,6%','GBD estimates','Ước tính mô hình']],
    widths=[2.5,2.0,2.5,3.0,3.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : ĐNA có tỷ lệ THẤP NHẤT (3,6%) \u2014 nhưng điều này có thể do GSHS chỉ dùng 1\u20132 câu hỏi lo âu (không phải thang chuẩn GAD-7 hay DASS-21). So với GAD-7 tại VN (Hoa 2024: 40,6%) hay DASS-21 (Bảo Quyên: 86,2%), sự chênh lệch 10\u201324 lần nhấn mạnh vai trò công cụ đo lường.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*ĐNA thấp nhất.* Có thể do (1) GSHS ít câu hỏi, (2) văn hóa Á Đông ít báo cáo lo âu, (3) hoặc thực sự thấp hơn. Cần NC xác nhận bằng GAD-7 trên cùng mẫu GSHS.')
bl(d,'*Bắt nạt và cô đơn.* Yếu tố nguy cơ nhất quán liên quốc gia, phù hợp Chen 2023 (bắt nạt OR = 1,51\u20131,97) và Ngo Anh Vinh 2024 (bạn bè kém OR = 6,84).')
bl(d,'*Hỗ trợ cha mẹ bảo vệ.* Phù hợp Pham 2024 VN (chăm sóc cảm xúc beta = \u20130,40) và Qiu 2022 (nuôi dạy tích cực OR = 0,30).')

rb(d,'Kết luận')
bl(d,'Dữ liệu GSHS từ 59 quốc gia, cho thấy ĐNA có tỷ lệ lo âu thấp nhất (3,6%) nhưng bắt nạt, cô đơn và thiếu hỗ trợ cha mẹ là yếu tố nguy cơ nhất quán liên quốc gia, gợi ý rằng can thiệp chống bắt nạt và tăng cường gắn kết gia đình có thể hiệu quả ở mọi bối cảnh văn hóa.',bold=True)

rh2(d,'Phản biện')
bl(d,'Q1 IF = 6,6, 59 nước, GSHS chuẩn hóa. Tuy nhiên, GSHS chỉ 1\u20132 câu lo âu (không phải GAD-7) \u2014 đánh giá thấp. Cắt ngang. Dữ liệu GSHS có thể cũ (thu thập nhiều năm).')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'So sánh GSHS vs GAD-7 trên cùng mẫu VTN ĐNA. Giải thích tại sao ĐNA thấp nhất. Phân tích riêng VN trong GSHS.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50 Cao. Q1, 59 nước, GSHS chuẩn hóa, nhưng công cụ đo hạn chế.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT31_59Countries.docx');sys.stderr.write('QT31 OK\n')

# ===== QT32: IRELAND MY WORLD (5 bảng gốc) =====
d=md()
bl(d,'Tóm tắt bài QT-32',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Khám phá xu hướng thay đổi trầm cảm và lo âu ở thanh thiếu niên từ 2012 đến 2019: Từ khảo sát My World cắt ngang lặp lại \u00bb (Exploring Changing Trends in Depression and Anxiety among Adolescents from 2012 to 2019: Insights from My World Repeated Cross-sectional Surveys), do Fitzgerald A và cộng sự (2024), School of Psychology, Dublin, Ireland, xuất bản trên Early Intervention in Psychiatry. 11 trang, 5 bảng.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Thiết kế cắt ngang lặp lại (repeated cross-sectional) — 2 đợt khảo sát 2012 và 2019 trên các mẫu VTN Ireland khác nhau. Nói cách khác, đây là thiết kế tốt hơn đơn cắt ngang vì cho phép đánh giá xu hướng thay đổi theo thời gian, mặc dù không theo dõi cùng cá nhân.',bold=True)
bl(d,'Các biến dự báo giải thích 37\u201361% phương sai trong trầm cảm và lo âu \u2014 tỷ lệ giải thích cao bất thường cho NC tâm lý.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 1-2 gốc). Xu hướng 2012 \u2192 2019:', bold=True)
tbl(d,['Chỉ số','2012','2019','Xu hướng'],
    [['Trầm cảm','Mức cơ sở','TĂNG có ý nghĩa','P < 0,05'],
     ['Lo âu','Mức cơ sở','TĂNG có ý nghĩa','P < 0,05'],
     ['Nữ','','Tăng NHANH hơn nam','Nhất quán y văn'],
     ['Nam','','Tăng chậm hơn','']],
    widths=[3.5,2.5,3.5,3.0])
d.add_paragraph()

bl(d,'Bảng 2 (từ Table 3-4 gốc). Yếu tố dự báo:', bold=True)
tbl(d,['Yếu tố','Loại','Nhất quán 2012 + 2019?'],
    [['Lòng tự trọng','Bảo vệ','Có \u2014 nhất quán cả 2 đợt'],
     ['Hỗ trợ xã hội','Bảo vệ','Có'],
     ['Hoạt động thể chất','Bảo vệ','Có'],
     ['Bắt nạt','Nguy cơ','Có'],
     ['Screen time / MXH','Nguy cơ','Có (tăng 2019)'],
     ['Stress','Nguy cơ','Có']],
    widths=[4.0,3.0,5.0])
d.add_paragraph()

bl(d,'Bảng 3. So sánh xu hướng Ireland với các nước:', bold=True)
tbl(d,['Nước','Giai đoạn','Xu hướng','Phương pháp'],
    [['Ireland (bài này)','2012\u20132019','Tăng, nữ nhanh hơn','Cắt ngang lặp lại'],
     ['Na Uy (QT21)','2011\u20132024','Tăng liên tục 13 năm','Decomposition'],
     ['Hoa Kỳ (QT23)','2013\u20132021','Lo âu +100% (AOR=2,17)','Hồ sơ lâm sàng'],
     ['Hàn Quốc (QT34)','2006\u20132022','Giảm trước, tăng sau COVID','KYRBS'],
     ['Việt Nam (VN14)','2021 vs 2023','41,5% \u2192 25,4%','DASS-21 2 thời điểm']],
    widths=[3.0,2.5,3.5,3.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : Xu hướng tăng lo âu/trầm cảm ở Ireland (2012\u20132019) nhất quán với Na Uy, Mỹ, Hàn Quốc. Nữ tăng nhanh hơn nam ở TẤT CẢ các nước. Các yếu tố bảo vệ (lòng tự trọng, hỗ trợ xã hội, thể chất) phù hợp với Pham 2024 VN (chăm sóc cảm xúc), Qiu 2022 (nuôi dạy tích cực), Zhu 2025 (hoạt động ngoài trời).')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*37\u201361% phương sai được giải thích.* Tỷ lệ rất cao cho NC tâm lý \u2014 gợi ý các yếu tố đã đo lường (tự trọng, hỗ trợ XH, bắt nạt, MXH) nắm bắt phần lớn cơ chế.')
bl(d,'*Nữ tăng nhanh hơn.* Nhất quán toàn cầu. Gợi ý can thiệp ưu tiên nữ VTN.')
bl(d,'*MXH tăng vai trò 2019 vs 2012.* Phù hợp với sự phổ biến smartphone + Nature 2025.')

rb(d,'Kết luận')
bl(d,'Dữ liệu cắt ngang lặp lại từ Ireland, cho thấy trầm cảm và lo âu ở VTN tăng đáng kể từ 2012 đến 2019 với nữ tăng nhanh hơn nam, và các yếu tố bảo vệ (tự trọng, hỗ trợ XH, thể chất) nhất quán cả 2 đợt, gợi ý rằng can thiệp tăng cường lòng tự trọng và hỗ trợ xã hội có thể hiệu quả xuyên thời gian và văn hóa.',bold=True)

rh2(d,'Phản biện')
bl(d,'Cắt ngang lặp lại \u2014 tốt hơn đơn cắt ngang. 37\u201361% phương sai giải thích \u2014 cao. Tuy nhiên, chỉ Ireland, 2 mẫu khác nhau (không dọc). Công cụ đo cụ thể cần kiểm tra từ PDF.')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'Cắt ngang lặp lại tương tự tại VN (2 đợt cách 5\u20137 năm). So sánh yếu tố dự báo Ireland vs VN. Can thiệp tăng lòng tự trọng tại trường THPT VN.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50 Cao. Cắt ngang lặp lại, xu hướng 7 năm, 37\u201361% phương sai, yếu tố nhất quán.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT32_Ireland.docx');sys.stderr.write('QT32 OK\n')
