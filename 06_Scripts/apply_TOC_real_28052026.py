# -*- coding: utf-8 -*-
"""Apply real heading pages vao bang TOC thu cong.
Khong dong vao ten muc (chi update page number).
28/05/2026."""
import os, sys, io, json, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
HEADINGS_FILE = os.path.join(ROOT, '06_Scripts', '_headings_real_28052026.json')

with open(HEADINGS_FILE, 'r', encoding='utf-8') as f:
    headings = json.load(f)
print(f"Loaded {len(headings)} headings")

# Normalize function: lowercase, strip whitespace + remove leading numbering
def normalize(s):
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    # Strip leading numbering like "1.", "1.1", "1.1.1.", "1.", etc.
    s = re.sub(r'^[\d\.]+\s*', '', s)
    s = s.strip(' .')
    return s


# Build lookup: normalized text -> page
lookup = {}
for h in headings:
    key = normalize(h['text'])
    if key and key not in lookup:
        lookup[key] = h['page']
    # Also full key (with numbering)
    full_key = h['text'].strip().lower()
    full_key = re.sub(r'\s+', ' ', full_key)
    if full_key and full_key not in lookup:
        lookup[full_key] = h['page']


# Open the LA
d = Document(FILE)
# Find TOC table
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb
        break

if toc_table is None:
    print("KHONG TIM THAY bang TOC")
    sys.exit(1)
print(f"TOC table: {len(toc_table.rows)} rows")

# Match each row
print()
print("=== MATCHING + UPDATE ===")
not_found = []
updated = 0
unchanged = 0
total = len(toc_table.rows)

for ri, row in enumerate(toc_table.rows):
    name = row.cells[0].text.strip()
    old_page = row.cells[1].text.strip() if len(row.cells) > 1 else ''
    norm = normalize(name)
    full = re.sub(r'\s+', ' ', name.strip().lower())
    # Try matches
    new_page = lookup.get(norm) or lookup.get(full)
    # Fuzzy: substring match
    if new_page is None and norm:
        for k, v in lookup.items():
            if k and (norm in k or (len(k) > 5 and k in norm)):
                # Strong containment match
                if len(k) > 5 or len(norm) > 5:
                    new_page = v
                    break
    if new_page is None:
        not_found.append((ri, name))
        continue
    new_page_str = str(new_page)
    if new_page_str != old_page:
        # Update cell
        cell = row.cells[1]
        # Get existing run formatting if possible
        font_name = 'Times New Roman'
        font_size = Pt(13)
        if cell.paragraphs and cell.paragraphs[0].runs:
            old_r = cell.paragraphs[0].runs[0]
            if old_r.font.name:
                font_name = old_r.font.name
            if old_r.font.size:
                font_size = old_r.font.size
        # Clear and add new
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''
        cell.paragraphs[0].text = ''
        new_run = cell.paragraphs[0].add_run(new_page_str)
        new_run.font.name = font_name
        new_run.font.size = font_size
        updated += 1
        if updated <= 20:
            print(f"  Row {ri:>2}: {name[:55]:<55} | {old_page:>4} -> {new_page_str}")
    else:
        unchanged += 1

print(f"\nUpdated: {updated}/{total}")
print(f"Unchanged (page already correct): {unchanged}")
print(f"Not found in document headings: {len(not_found)}")
if not_found:
    print("\nDanh sach KHONG TIM THAY (giu nguyen page cu):")
    for ri, name in not_found[:30]:
        print(f"  Row {ri:>2}: {name}")

# Backup
backup = FILE.replace('.docx', '_BEFORE_TOC_PAGE_UPDATE.docx')
shutil.copy(FILE, backup)
print(f"\nBackup: {backup}")
d.save(FILE)
print(f"Saved: {FILE}")
print(f"Size: {os.path.getsize(FILE)//1024} KB")
