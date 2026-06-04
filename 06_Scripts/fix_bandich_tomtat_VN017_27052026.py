# -*- coding: utf-8 -*-
"""Fix ban dich + tom tat VN017 Nguyen Danh Lam 2022 theo so lieu PDF goc.
Cac loi can sua:
1. Bang 2 (Stress): Nhe 16,0%->18,7%; Vua 14,1%->15,1%; Nang 6,8%->7,1%; Rat nang 4,8%->0,8%
2. Bang 2 (Lo au): Nhe 7,7%->11,2%; Vua 24,5%->25,1%
3. "ban do thi" -> "ban nong thon" (nhieu cho)
4. Bo claim "tieu de ghi 2024"
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# 2 files can sua
FILES = [
    ('03_Ban-dich/VN017_NguyenDanhLam_2022_YHVN.docx', 'BAN DICH VN017'),
    ('Tom-tat-tung-bai/VN017_NguyenDanhLam_2022_YHVN.docx', 'TOM TAT VN017'),
]

# Map sua chi tiet so lieu trong cell bang
CELL_FIXES = {
    # Stress row
    'Nhe Stress old': ('16,0%', '18,7%'),
    'Vua Stress old': ('14,1%', '15,1%'),
    'Nang Stress old': ('6,8%', '7,1%'),
    'Rat nang Stress old': ('4,8%', '0,8%'),
    # Lo au row
    'Nhe Lo au old': ('7,7%', '11,2%'),
    'Vua Lo au old': ('24,5%', '25,1%'),
}

# Text sua trong paragraph
TEXT_REPLACEMENTS = [
    # Sua "ban do thi" -> "ban nong thon"
    ('Vùng bán đô thị Thanh Hóa', 'Huyện bán nông thôn Thanh Hóa'),
    ('vùng bán đô thị Thanh Hóa', 'huyện bán nông thôn Thanh Hóa'),
    ('vùng bán đô thị', 'huyện bán nông thôn'),
    # Bo claim "tieu de ghi 2024"
    ('Tạp chí không lập chỉ mục quốc tế. Năm xuất bản 2022 nhưng tiêu đề ghi 2024.',
     'Tạp chí không lập chỉ mục quốc tế.'),
]


def apply_text_fixes_paragraph(p, applied_list, file_label):
    """Fix text trong paragraph, mark sua mau DO."""
    full = p.text
    changed = False
    for old, new in TEXT_REPLACEMENTS:
        if old in full:
            # Track applied
            applied_list.append(f'[{file_label} para] "{old[:50]}..." -> "{new[:50]}..."')
            # Find positions of all occurrences
            positions = []
            cursor = 0
            while True:
                idx = full.find(old, cursor)
                if idx == -1: break
                positions.append((idx, idx + len(old)))
                cursor = idx + len(old)
            if not positions:
                continue
            # Build segments
            segments = []
            cur = 0
            for s, e in positions:
                if cur < s:
                    segments.append((full[cur:s], False))
                segments.append((new, True))
                cur = e
            if cur < len(full):
                segments.append((full[cur:], False))
            # Rebuild
            for r in p.runs:
                r.text = ''
            for text, is_red in segments:
                if not text: continue
                r = p.add_run(text)
                r.font.name = 'Times New Roman'; r.font.size = Pt(13)
                if is_red:
                    r.font.color.rgb = RGBColor(192, 0, 0)
                    r.bold = True
            full = ' '.join(s[0] for s in segments)
            changed = True
    return changed


def fix_table_cells(d, applied_list, file_label):
    """Find Stress va Lo au rows in tables, fix specific cells."""
    for ti, tb in enumerate(d.tables):
        # Iterate rows
        rows = tb.rows
        for ri, row in enumerate(rows):
            first_cell_text = row.cells[0].text.strip()
            if first_cell_text == 'Stress' and len(row.cells) >= 6:
                # Stress row: cells = [Stress, Tỷ lệ, Nhẹ, Vừa, Nặng, Rất nặng]
                fixes = [
                    (2, '16,0%', '18,7%', 'Stress Nhẹ'),
                    (3, '14,1%', '15,1%', 'Stress Vừa'),
                    (4, '6,8%', '7,1%', 'Stress Nặng'),
                    (5, '4,8%', '0,8%', 'Stress Rất nặng'),
                ]
                for ci, old_val, new_val, label in fixes:
                    cell = row.cells[ci]
                    if old_val in cell.text:
                        # Replace text - clear and add red run
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = ''
                        cell.paragraphs[0].text = ''
                        new_run = cell.paragraphs[0].add_run(new_val)
                        new_run.font.name = 'Times New Roman'
                        new_run.font.size = Pt(13)
                        new_run.font.color.rgb = RGBColor(192, 0, 0)
                        new_run.bold = True
                        applied_list.append(f'[{file_label} Table {ti+1} {label}] {old_val} -> {new_val}')
            elif first_cell_text == 'Lo âu' and len(row.cells) >= 6:
                fixes = [
                    (2, '7,7%', '11,2%', 'Lo âu Nhẹ'),
                    (3, '24,5%', '25,1%', 'Lo âu Vừa'),
                ]
                for ci, old_val, new_val, label in fixes:
                    cell = row.cells[ci]
                    if old_val in cell.text:
                        for p in cell.paragraphs:
                            for r in p.runs:
                                r.text = ''
                        cell.paragraphs[0].text = ''
                        new_run = cell.paragraphs[0].add_run(new_val)
                        new_run.font.name = 'Times New Roman'
                        new_run.font.size = Pt(13)
                        new_run.font.color.rgb = RGBColor(192, 0, 0)
                        new_run.bold = True
                        applied_list.append(f'[{file_label} Table {ti+1} {label}] {old_val} -> {new_val}')


# Apply fixes to both files
for src_relpath, label in FILES:
    src_full = os.path.join(ROOT, src_relpath)
    # Create output filename with _fixed suffix
    base, ext = os.path.splitext(src_relpath)
    out_relpath = base + '_FIXED_27052026' + ext
    out_full = os.path.join(ROOT, out_relpath)
    print(f"\n=== {label} ===")
    print(f"Source: {src_full}")
    print(f"Output: {out_full}")

    d = Document(src_full)
    applied = []

    # Fix text in paragraphs
    for p in d.paragraphs:
        apply_text_fixes_paragraph(p, applied, label)

    # Fix cells in tables
    fix_table_cells(d, applied, label)

    d.save(out_full)
    print(f"Applied {len(applied)} fixes:")
    for a in applied:
        print(f"  ✓ {a}")
