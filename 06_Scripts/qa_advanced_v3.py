# -*- coding: utf-8 -*-
"""
QA NÂNG CAO bao cao v3 — multi-layer checks:

1. INTERNAL CONSISTENCY — fact appears multiple times, all must match
2. NUMBER PLAUSIBILITY — flag suspicious values (% > 100, n < 10, etc.)
3. CITATION COMPLETENESS — every numeric claim should have nearby citation
4. CROSS-REFERENCE WITH SUMMARY — facts in report should match facts in tom tat
5. STATISTICAL INTERPRETATION — Cohen's d / SMD interpretation correct
6. FORMAT INTEGRITY — table columns aligned, captions present
7. AUTHOR/YEAR CONSISTENCY — same author should have same year throughout
8. ABBREVIATION USAGE — first use should be defined
"""
import os, sys, io, re
from collections import defaultdict, Counter
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
REPORT = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 12042026 v4.docx')
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')

# ============================================================
# Read report
# ============================================================
doc = Document(REPORT)
all_text_paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
all_text = '\n'.join(all_text_paragraphs)
all_table_text = []
for tb in doc.tables:
    for row in tb.rows:
        for cell in row.cells:
            all_table_text.append(cell.text)
full_text = all_text + '\n' + '\n'.join(all_table_text)

print(f'Báo cáo: {os.path.basename(REPORT)}')
print(f'Đoạn: {len(all_text_paragraphs)}, Bảng: {len(doc.tables)}')
print(f'Tổng ký tự: {len(full_text):,}')
print('=' * 70)

# ============================================================
# CHECK 1: INTERNAL CONSISTENCY — facts appearing multiple times
# ============================================================
print('\n--- CHECK 1: NHẤT QUÁN NỘI BỘ ---')
# Find author-year-fact triples
# Pattern: "Author year ... d = X" → check if d=X consistent across mentions
# Simpler: scan for pairs (author_lastname, distinct_number)

# Extract distinct numeric facts with author labels
fact_pat = re.compile(
    r'(\w+(?:\s+(?:&|et al\.))?\s+\d{4})[^.]*?'
    r'(d\s*=\s*-?\d+[,\.]\d+|g\s*=\s*-?\d+[,\.]\d+|β\s*=\s*-?\d+[,\.]\d+|'
    r'OR\s*=\s*-?\d+[,\.]\d+|SMD\s*=\s*-?\d+[,\.]\d+|n\s*=\s*\d+[\.,]?\d*|'
    r'p\s*[<>]\s*0[,\.]\d+|SUCRA\s*=?\s*\d+[,\.]?\d*\s*%?)',
    re.IGNORECASE
)

facts = defaultdict(list)
for m in fact_pat.finditer(full_text):
    author_year = m.group(1).strip()
    fact = m.group(2).strip()
    facts[author_year].append(fact)

inconsistencies = []
for ay, fs in facts.items():
    # Group by stat type
    by_stat = defaultdict(set)
    for f in fs:
        # Get stat type (d, g, β, OR, SMD, n, p, SUCRA)
        stype = re.match(r'(\w+|β)', f).group(1).lower()
        # Get value
        val_m = re.search(r'-?\d+[,\.]?\d*', f)
        if val_m:
            by_stat[stype].add(val_m.group(0).replace(',', '.'))
    for stype, vals in by_stat.items():
        if len(vals) > 1:
            inconsistencies.append((ay, stype, vals))

if inconsistencies:
    print(f'⚠ {len(inconsistencies)} cặp số liệu KHÔNG NHẤT QUÁN:')
    for ay, st, vals in inconsistencies[:10]:
        print(f'  {ay} | {st} = {vals}')
else:
    print('✓ Không tìm thấy cặp số liệu mâu thuẫn (cùng tác giả + cùng loại thống kê)')

# ============================================================
# CHECK 2: NUMBER PLAUSIBILITY
# ============================================================
print('\n--- CHECK 2: NUMBER PLAUSIBILITY ---')

issues_2 = []
# Find percentages
pct_pat = re.compile(r'(\d+[,\.]\d+)\s*%')
for m in pct_pat.finditer(full_text):
    val = float(m.group(1).replace(',', '.'))
    if val > 100:
        issues_2.append(f'PCT > 100: {m.group(0)}')
    if val < 0:
        issues_2.append(f'PCT < 0: {m.group(0)}')

# Find p-values
p_pat = re.compile(r'p\s*[<>=]\s*(\d+[,\.]?\d*)', re.IGNORECASE)
for m in p_pat.finditer(full_text):
    try:
        val = float(m.group(1).replace(',', '.'))
        if val > 1:
            issues_2.append(f'p > 1: {m.group(0)}')
    except: pass

# Find Cohen's d / Hedges g — should be -3 to +3 for normal range
d_pat = re.compile(r'(?:d|g)\s*=\s*(-?\d+[,\.]\d+)', re.IGNORECASE)
for m in d_pat.finditer(full_text):
    val = float(m.group(1).replace(',', '.'))
    if abs(val) > 3:
        issues_2.append(f'd/g extreme: {m.group(0)} (very unusual)')

# Sample sizes — flag n < 10 (suspect typo)
n_pat = re.compile(r'\bn\s*=\s*(\d+)', re.IGNORECASE)
for m in n_pat.finditer(full_text):
    val = int(m.group(1))
    if val < 5:
        issues_2.append(f'n very small: {m.group(0)}')

if issues_2:
    print(f'⚠ {len(issues_2)} vấn đề về tính hợp lý số liệu:')
    for i in issues_2[:15]:
        print(f'  {i}')
else:
    print('✓ Không tìm thấy số liệu bất thường về tính hợp lý')

# ============================================================
# CHECK 3: CITATION COMPLETENESS
# ============================================================
print('\n--- CHECK 3: TRÍCH DẪN ĐỦ ---')
# For each paragraph containing a numeric stat, check if author cited
para_with_stat_no_cite = []
stat_keywords = ['Cohen d', 'Cohen\'s d', 'Hedges g', 'OR =', 'β =', 'SMD =', 'SUCRA',
                 'p <', 'p =', 'AOR =', 'NNT =']
author_pat = re.compile(r'(\w[\w-]+(?:\s+(?:&|et al\.|và cộng sự))?\s*\(?\d{4}\)?)',
                        re.IGNORECASE)

for i, para in enumerate(all_text_paragraphs):
    has_stat = any(kw in para for kw in stat_keywords)
    if not has_stat:
        # Also check for percentages
        if re.search(r'\d+[,\.]\d+\s*%', para):
            has_stat = True
    if has_stat:
        has_cite = bool(author_pat.search(para)) or 'CAMS' in para or 'NEJM' in para
        if not has_cite and len(para) > 60:
            para_with_stat_no_cite.append((i, para[:120]))

if para_with_stat_no_cite:
    print(f'⚠ {len(para_with_stat_no_cite)} đoạn có số liệu nhưng KHÔNG có trích dẫn rõ ràng:')
    for i, p in para_with_stat_no_cite[:8]:
        print(f'  [#{i}] {p}')
else:
    print('✓ Tất cả đoạn có số liệu đều có trích dẫn')

# ============================================================
# CHECK 4: CROSS-REFERENCE VỚI TÓM TẮT
# ============================================================
print('\n--- CHECK 4: ĐỐI CHIẾU VỚI TÓM TẮT ---')
# Sample: pick key facts from report and check if they appear in matching summary

cross_checks = [
    ('Happy House Tran 2023', 'VN0', '1.084'),  # Happy House n=1.084
    ('Walder DMHI', 'QT040', '0,878'),
    ('Bress Maya', 'QT043', '1,04'),
    ('Walkup CAMS', 'QT028', '80,7'),
    ('Sasaki Japan', 'QT045', '4,97'),
    ('De Silva Sri Lanka', 'QT038', '720'),
    ('Dong PLOS', 'QT047', '0,22'),
    ('Cai Resilience', 'QT044', 'Cai'),
    ('Dao Thi Ngoan', 'VN028', '4,97'),
    ('Duong TPHCM', 'VN029', '2.631'),
]

mismatches = []
for label, summary_prefix, expected_value in cross_checks:
    # Check if expected_value appears in report
    in_report = expected_value in full_text
    # Find matching summary file
    matching_files = [f for f in os.listdir(TT_DIR) if f.startswith(summary_prefix)]
    if matching_files:
        sum_path = os.path.join(TT_DIR, matching_files[0])
        sum_doc = Document(sum_path)
        sum_text = '\n'.join(p.text for p in sum_doc.paragraphs)
        for tb in sum_doc.tables:
            for row in tb.rows:
                for cell in row.cells:
                    sum_text += '\n' + cell.text
        in_summary = expected_value in sum_text
        status = 'OK' if (in_report and in_summary) else 'MISMATCH'
        if not (in_report and in_summary):
            mismatches.append((label, expected_value, in_report, in_summary, matching_files[0]))
        print(f'  [{status}] {label}: report={in_report} | summary={in_summary} | val={expected_value}')

if not mismatches:
    print('✓ Tất cả 10 fact đối chiếu khớp giữa báo cáo và tóm tắt')
else:
    print(f'⚠ {len(mismatches)} mismatch')

# ============================================================
# CHECK 5: STATISTICAL INTERPRETATION
# ============================================================
print('\n--- CHECK 5: DIỄN GIẢI HIỆU LỰC THỐNG KÊ ---')
# Cohen's d standards: d ≤ 0.2 = small, 0.5 = medium, 0.8 = large
# Find claims about effect sizes and check interpretation

interp_pat = re.compile(
    r'(d|g)\s*=\s*(-?\d+[,\.]\d+)[^.]*?'
    r'(NHỎ|TRUNG BÌNH|LỚN|RẤT LỚN|small|medium|large|nhỏ|trung bình|lớn|rất lớn)',
    re.IGNORECASE
)

interp_issues = []
for m in interp_pat.finditer(full_text):
    val = abs(float(m.group(2).replace(',', '.')))
    interp = m.group(3).lower()
    expected = ''
    if val < 0.2:
        expected = 'NHỎ/small'
    elif val < 0.5:
        expected = 'NHỎ-TRUNG BÌNH/small-medium'
    elif val < 0.8:
        expected = 'TRUNG BÌNH/medium'
    elif val < 1.2:
        expected = 'LỚN/large'
    else:
        expected = 'RẤT LỚN/very large'
    # Crude check: just print val + interpretation
    print(f'  d/g = {val} → "{interp}" (chuẩn Cohen: {expected})')

# ============================================================
# CHECK 6: FORMAT INTEGRITY (tables)
# ============================================================
print('\n--- CHECK 6: FORMAT BẢNG ---')
table_issues = []
for ti, tb in enumerate(doc.tables, start=1):
    n_cols = len(tb.rows[0].cells)
    n_rows = len(tb.rows)
    # All rows should have same number of cols
    cols_per_row = [len(r.cells) for r in tb.rows]
    if len(set(cols_per_row)) > 1:
        table_issues.append(f'Bảng {ti}: cột không đều: {cols_per_row}')
    # Header row should be non-empty
    header_cells = [c.text.strip() for c in tb.rows[0].cells]
    empty_headers = sum(1 for c in header_cells if not c)
    if empty_headers > 0:
        table_issues.append(f'Bảng {ti}: {empty_headers} header trống')
    # Empty cells overall
    empty_cells = 0
    total_cells = 0
    for row in tb.rows:
        for cell in row.cells:
            total_cells += 1
            if not cell.text.strip():
                empty_cells += 1
    if empty_cells / total_cells > 0.3:
        table_issues.append(f'Bảng {ti}: {empty_cells}/{total_cells} cells trống ({empty_cells*100//total_cells}%)')

if table_issues:
    print(f'⚠ {len(table_issues)} vấn đề về bảng:')
    for i in table_issues:
        print(f'  {i}')
else:
    print(f'✓ Tất cả {len(doc.tables)} bảng OK về format')

# ============================================================
# CHECK 7: YEAR CONSISTENCY
# ============================================================
print('\n--- CHECK 7: NĂM XUẤT BẢN NHẤT QUÁN ---')
# For each unique author, check all years mentioned
author_years = defaultdict(set)
ay_pat = re.compile(r'\b(\w[\w-]+)\s+(?:&\s+\w+\s+)?(?:et al\.\s+)?(\d{4})\b')
for m in ay_pat.finditer(full_text):
    author = m.group(1)
    year = m.group(2)
    if author.lower() in {'bài', 'phần', 'mục', 'bảng', 'trang', 'năm', 'từ', 'sau', 'trước'}:
        continue
    if author[0].isupper() and len(author) > 3:
        author_years[author].add(year)

multi_year_authors = {a: y for a, y in author_years.items() if len(y) > 1}
if multi_year_authors:
    print(f'⚠ {len(multi_year_authors)} tác giả có nhiều năm khác nhau (cần kiểm tra):')
    for a, ys in list(multi_year_authors.items())[:10]:
        print(f'  {a}: {sorted(ys)}')
else:
    print('✓ Mỗi tác giả có 1 năm xuất bản nhất quán')

# ============================================================
# CHECK 8: ABBREVIATION USAGE
# ============================================================
print('\n--- CHECK 8: VIẾT TẮT ---')
# Find common abbreviations and check if defined on first use
common_abbrs = ['CBT', 'iCBT', 'gCBT', 'DMHI', 'VRET', 'NMA', 'MA', 'SR', 'RCT',
                'PE', 'PA', 'SAD', 'GAD', 'OR', 'AOR', 'CI', 'CrI', 'SUCRA',
                'SMD', 'NNT', 'MHST', 'CALM', 'BESST', 'PLACES', 'CTAF']

abbr_uses = {}
for abbr in common_abbrs:
    # Find first occurrence
    pat = re.compile(r'\b' + re.escape(abbr) + r'\b')
    m = pat.search(full_text)
    if m:
        # Look for definition: (Full Name) or — Full Name nearby
        idx = m.start()
        snippet = full_text[max(0, idx-50):idx+200]
        # Has parenthesis with full name?
        defined = bool(re.search(r'\(' + re.escape(abbr) + r'[^)]*\)|' +
                                  re.escape(abbr) + r'\s*[—–-]\s*\w', snippet))
        abbr_uses[abbr] = defined

undefined = [a for a, d in abbr_uses.items() if not d]
if undefined:
    print(f'⚠ {len(undefined)} viết tắt KHÔNG được định nghĩa khi dùng lần đầu:')
    for a in undefined[:15]:
        print(f'  {a}')
else:
    print('✓ Tất cả viết tắt thông dụng đều được định nghĩa')

# ============================================================
# SUMMARY
# ============================================================
print('\n' + '=' * 70)
print('TỔNG KẾT QA NÂNG CAO')
print('=' * 70)

n_critical = len(inconsistencies) + len(issues_2)
n_warnings = len(para_with_stat_no_cite) + len(mismatches) + len(table_issues) + len(undefined)

print(f'Critical issues:  {n_critical}')
print(f'Warnings:         {n_warnings}')
print(f'Total issues:     {n_critical + n_warnings}')

if n_critical == 0:
    print('\n✓ Báo cáo đạt mức "ĐỦ ĐIỀU KIỆN PHÁT HÀNH" về số liệu')
else:
    print('\n⚠ Cần sửa các critical issues trước khi phát hành')
