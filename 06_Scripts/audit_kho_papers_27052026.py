# -*- coding: utf-8 -*-
"""Audit chi tiet cac bai trong kho - cross-check LA claims vs tom-tat content.
27/05/2026.

Cac bai uu tien check:
- Mandaknalli & Malusare 2021
- Chen 2023 (Trung Quoc)
- Wen 2020
- Qiu 2022
- Anderson 2025
- Zhu 2025
- Saikia 2023 (do them - voi PDF da audit)
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_4_FixAbbrev_27052026.docx')
d = Document(LA)

# Author -> file in Tom-tat-tung-bai
AUTHORS_TO_CHECK = [
    ('Mandaknalli', 'QT003_Mandaknalli_Malusare_2021.docx'),
    ('Chen', 'QT007_Chen_et_al_2023_China_BMCPsych.docx'),
    ('Wen', 'QT008_Wen_et_al_2020_China_Rural.docx'),
    ('Qiu', 'QT009_Qiu_et_al_2022_China_Parenting.docx'),
    ('Anderson', 'QT014_Anderson_2025_Wiley_Narrative.docx'),
    ('Zhu', 'QT015_Zhu_2025_BMC_China.docx'),
    ('Saikia', 'QT002_Saikia_et_al_2023_India_Assam.docx'),
    ('Jenkins', 'QT001_Jenkins_et_al_2023_USA.docx'),
]

def extract_numbers(text):
    """Extract all numeric values with %."""
    pattern = re.compile(r'(\d+[\.,]\d+|\d+)\s*%')
    return pattern.findall(text)


for author, tt_file in AUTHORS_TO_CHECK:
    print(f"\n{'=' * 70}")
    print(f"=== {author} ===")
    print(f"{'=' * 70}")
    # LA paras
    la_paras = [(i, p.text) for i, p in enumerate(d.paragraphs)
                if author in p.text and i < 1300]
    print(f"\nLA paras with '{author}': {len(la_paras)}")
    # Print short context with numbers
    for i, txt in la_paras[:5]:
        nums = extract_numbers(txt)
        # find sentence with author + nearby numbers
        pos = txt.find(author)
        ctx = txt[max(0, pos-80):min(len(txt), pos+300)]
        print(f"  para {i}: ...{ctx}...")
        if nums:
            print(f"    [nums in context: {nums[:8]}]")
    # Tom-tat content
    tt_path = os.path.join(ROOT, 'Tom-tat-tung-bai', tt_file)
    if not os.path.exists(tt_path):
        print(f"\n  !!! Tom-tat KHONG TON TAI: {tt_file}")
        continue
    td = Document(tt_path)
    tt_text = '\n'.join(p.text for p in td.paragraphs)
    tt_nums = extract_numbers(tt_text)
    print(f"\n  Tom-tat numbers: {tt_nums[:30]}")
    # Print n=, N=, sample size keywords
    for kw in ['n =', 'n=', 'N =', 'N=', 'mẫu', 'sample', 'students', 'học sinh']:
        for i, p in enumerate(td.paragraphs):
            if kw in p.text and ('học sinh' in p.text or 'student' in p.text.lower() or 'mẫu' in p.text or any(c.isdigit() for c in p.text[:200])):
                snippet = p.text[:300]
                if any(c.isdigit() for c in snippet):
                    print(f"    [tt para {i} ({kw})]: {snippet}")
                    break
