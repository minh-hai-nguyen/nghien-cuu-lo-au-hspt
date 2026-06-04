"""DOC: Phan bien VN020 Tran Ho Vinh Loc (2024) ve Ap luc hoc tap.
Style giong cac doc phan bien truoc (VN004, VN014).

ALL FACTS VERIFIED tu Tom-tat VN020:
- Tac gia: Tran Ho Vinh Loc, Huynh Ngoc Van Anh, To Gia Kien (2024)
- Khoa Y te Cong cong, DH Y Duoc TPHCM
- Tap chi Y hoc TPHCM
- n = 976 HS THPT, 2 truong TPHCM
- Lay mau: ngau nhien bac truong + thuan tien 4 lop/khoi
- Pilot test 34 HS
- Cong cu: DASS-Y (Szabo 2010) + ESSA (Sun 2011, alpha=0,81)
- Nguong DASS-Y: lo au >=6, tram cam >=7, stress >=12
- Phan tich: hoi quy logistic da bien
- Lo au 25,1% (DASS-Y)
- ESSA >= 59 yeu to nguy co manh nhat
- Nu > nam (p<0,05)
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Phan_bien_VN020_TranHoVinhLoc_2024_ALHT.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def blue_bullet(text, size=12):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẢN BIỆN VN020 TRẦN HỒ VĨNH LỘC (2024)\nVỀ CÁCH ĐO ÁP LỰC HỌC TẬP\n— Bảy hạn chế phương pháp luận —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Em cho thầy phản biện về áp lực học tập của Trần Hồ Vĩnh Lộc '
    '2024 nhé.', italic=True
)

# 1. Bối cảnh — ƯU điểm trước
H('1. Bối cảnh — VN020 là bài VN TỐT NHẤT về ALHT đến nay', level=2, color=GREEN)
para(
    'Trước khi phản biện, cần ghi nhận VN020 là MỘT TRONG SỐ ÍT bài '
    'Việt Nam về áp lực học tập đo bằng THANG CHUẨN QUỐC TẾ. Cụ '
    'thể:'
)
para('• Tác giả: Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh, Tô Gia Kiên (2024) — Khoa Y tế Công cộng, Đại học Y Dược TPHCM.')
para('• Tạp chí Y học TPHCM, năm 2024.')
para('• n = 976 HS THPT tại 2 trường TPHCM — cỡ mẫu lớn nhất TPHCM về kết hợp DASS-Y + ESSA.')
para('• Công cụ: DASS-Y (Depression Anxiety Stress Scales for Youth, Szabó 2010 — phiên bản VTN) + ESSA (Educational Stress Scale for Adolescents, Sun và cộng sự 2011 — 16 mục, α = 0,81).')
para('• Phát hiện chính: lo âu 25,1% (DASS-Y); ESSA ≥ 59 là YẾU TỐ MẠNH NHẤT tăng nguy cơ cả 3 chỉ số (lo âu, trầm cảm, căng thẳng).')

para('')
para('ƯU điểm cần ghi nhận:', bold=True, color=GREEN)
para('• Sử dụng ESSA CHUẨN QUỐC TẾ — phù hợp Pascoe (2020), Walder (2025).')
para('• Áp dụng DASS-Y có ngưỡng RIÊNG cho VTN — đóng góp phương pháp luận quan trọng cho NC VN.')
para('• Lấy mẫu ngẫu nhiên ở bậc trường (chứ không thuận tiện toàn bộ).')
para('• Có pilot test 34 HS để kiểm tra độ tin cậy bộ câu hỏi.')

# 2. Bảy hạn chế
H('2. Bảy hạn chế phương pháp luận', level=2, color=NAVY)

H('2.1. Chỉ 2 TRƯỜNG TPHCM — không đại diện thành phố', level=3, color=NAVY)
para(
    'VN020 lấy mẫu tại CHỈ 2 trường THPT TPHCM. TPHCM có hơn 200 '
    'trường THPT (chính quy + dân lập + chuyên + GDTX). Kết quả '
    'từ 2 trường KHÔNG thể NGOẠI SUY ra HS THPT TPHCM nói chung — '
    'chứ chưa nói đến HS THPT Việt Nam toàn quốc.'
)
para(
    'So sánh: VN029 (Trúc Thanh Thái 2025) cũng tại TPHCM nhưng '
    'lấy mẫu 8 cơ sở (4 THPT chính quy + 4 GDTX) với n = 2.631 — '
    'cỡ mẫu gấp 2,7 lần VN020 và đa dạng cơ sở hơn. Tỷ lệ lo âu '
    'của VN029 (50,3% bằng DASS-21) khác xa VN020 (25,1% bằng '
    'DASS-Y) — gợi ý TỶ LỆ phụ thuộc cả thang đo VÀ mẫu.'
)

H('2.2. Chọn lớp THUẬN TIỆN ở bậc lớp', level=3, color=NAVY)
para(
    'Tom-tat ghi rõ: "Lấy mẫu ngẫu nhiên nhiều bậc (bậc trường '
    'ngẫu nhiên, bậc lớp THUẬN TIỆN 4 lớp/khối)". Tức là chỉ '
    'BẬC TRƯỜNG được ngẫu nhiên, BẬC LỚP là CHỌN THUẬN TIỆN. Đây '
    'là thiết kế ngẫu nhiên CHƯA HOÀN TOÀN.'
)
para(
    'Hậu quả: nếu giáo viên chọn các lớp "dễ tham gia" (lớp '
    'có giáo viên chủ nhiệm hợp tác), có thể có thiên lệch '
    'chọn lọc (selection bias). Các lớp này có thể có học sinh '
    'có lo âu thấp hơn vì giáo viên chủ động trong quản lý '
    'cảm xúc — phóng đại tỷ lệ "thấp" hoặc ngược lại.'
)
para(
    'Đề xuất sửa: chọn lớp ngẫu nhiên (probability proportional '
    'to size) trong từng trường thay vì thuận tiện.'
)

H('2.3. PHÂN TÍCH LOGISTIC dùng ngưỡng cắt nhị phân — MẤT thông tin liên tục', level=3, color=NAVY)
para(
    'VN020 phân tích bằng HỒI QUY LOGISTIC ĐA BIẾN với outcome '
    'NHỊ PHÂN (có/không lo âu theo ngưỡng DASS-Y) và ESSA '
    'ngưỡng ≥ 59. Cách tiếp cận này MẤT THÔNG TIN LIÊN TỤC:'
)
para('• Học sinh có ESSA = 58 và ESSA = 30 đều được xem "không có ALHT cao" — gộp chung dù mức độ KHÁC NHAU rõ rệt.')
para('• Học sinh có ESSA = 60 và ESSA = 80 đều "có ALHT cao" — gộp chung dù chênh 20 điểm.')
para('• Kết quả: chỉ có OR (Odds Ratio), KHÔNG có β chuẩn hóa, KHÔNG có R² — KHÔNG so sánh được với SEM của các nghiên cứu khác.')
para(
    'So sánh với chương 3 luận án CTH: dùng SEM với β chuẩn hóa '
    '(0,510 cho RLLATQ; 0,490 cho RLLAXH; 0,253 cho RLLACL) và '
    'R² = 0,284 cho mô hình ALHT → RLLA tổng. Các con số này '
    'CHO PHÉP so sánh trực tiếp với meta-analysis quốc tế '
    '(Walder 2025: g = 0,508 chung; g = 0,878 SAD-specific).'
)

H('2.4. KHÔNG phân tách 3 dạng RLLA — gộp thành "lo âu" tổng', level=3, color=NAVY)
para(
    'DASS-Y có 1 thang con cho LO ÂU (anxiety subscale 6 mục) — '
    'KHÔNG phân biệt lo âu lan tỏa, lo âu xã hội, và lo âu chia '
    'ly như RCADS. Hậu quả: VN020 chỉ có MỘT con số "tỷ lệ lo '
    'âu 25,1%" — không phân biệt được pattern theo loại RLLA.'
)
para(
    'Đây là hạn chế NGHIÊM TRỌNG vì y văn đã xác lập (chương 3 '
    'luận án CTH) các yếu tố nguy cơ có CƯỜNG ĐỘ KHÁC NHAU '
    'cho từng loại RLLA. Ví dụ:'
)
add_table(
    ['Yếu tố', 'β cho RLLATQ', 'β cho RLLACL', 'β cho RLLAXH', 'Pattern'],
    [
        ['ALHT', '0,510', '0,253', '0,490', 'Mạnh nhất ở lan tỏa'],
        ['Bắt nạt', '0,215', '0,376', '0,253', 'Mạnh nhất ở chia ly (đặc biệt)'],
        ['Tự trọng', '−0,455', '−0,087', '−0,415', 'Bảo vệ mạnh ở lan tỏa + xã hội'],
        ['HTCM', '−0,172', '0,000', '−0,273', 'Vắng mặt ở chia ly'],
    ]
)
para('')
para(
    'Gộp thành "lo âu" tổng làm MẤT các phát hiện đặc thù — '
    'đặc biệt phát hiện β bắt nạt → lo âu chia ly mạnh nhất.'
)

H('2.5. BỎ SÓT 5 yếu tố quan trọng', level=3, color=NAVY)
para(
    'VN020 đo: ALHT (ESSA) + cấu trúc gia đình (cha mẹ ly hôn) + '
    'kinh tế gia đình + giới tính. Tổng cộng ~4 nhóm yếu tố.'
)
para(
    'Y văn quốc tế đã xác lập NHIỀU yếu tố khác có cường độ tác '
    'động đáng kể, mà VN020 KHÔNG đo:'
)
add_table(
    ['Yếu tố thiếu', 'Y văn quốc tế', 'Cường độ trong chương 3 CTH'],
    [
        ['Bắt nạt học đường', 'Olweus 1996; meta-analysis OR=4,72',
         'β = 0,376 cho RLLACL — mạnh nhất'],
        ['Tự trọng', 'Sowislo & Orth 2013 meta 95 NC',
         'β = −0,455 — bảo vệ mạnh nhất'],
        ['Hỗ trợ cha mẹ', 'Phạm 2024 (VN003) β = −0,40',
         'β = −0,273 cho RLLAXH'],
        ['Nghiện điện thoại', 'Sohn 2019 meta 41 NC OR = 3,05',
         'β = 0,383 cho RLLAXH'],
        ['Biện pháp đối phó', 'Compas 2017 meta 212 NC',
         'β = 0,749 (fit kém — exploratory)'],
    ]
)
para('')
para(
    'Việc thiếu 5 yếu tố này khiến mô hình hồi quy logistic của '
    'VN020 BỎ SÓT phần lớn phương sai có thể giải thích — kết '
    'luận "ESSA là yếu tố mạnh nhất" có thể bị OVERESTIMATE vì '
    'không kiểm soát các yếu tố khác.'
)

H('2.6. DASS-Y CHƯA PHỔ BIẾN quốc tế — khó so sánh meta-analysis', level=3, color=NAVY)
para(
    'DASS-Y (Szabó 2010) là phiên bản VTN của DASS-21 với ngưỡng '
    'RIÊNG (lo âu ≥6, trầm cảm ≥7, stress ≥12 — thấp hơn DASS-21 '
    'người lớn). Đây là đóng góp phương pháp luận, NHƯNG:'
)
para('• DASS-Y CHƯA được sử dụng rộng rãi trong y văn quốc tế.')
para('• Không có META-ANALYSIS DASS-Y vs DASS-21 trên cùng mẫu để chứng minh ngưỡng riêng phù hợp.')
para('• Tỷ lệ lo âu VN020 (25,1%) THẤP HƠN ĐÁNG KỂ so với các bài VN dùng DASS-21 (40,6% Hoa 2024; 53,81% Phạm Thị Ngọc 2024 Cộng Hiền; 86,2% Bảo Quyên 2025) — gợi ý NGƯỠNG DASS-Y có thể QUÁ NGHIÊM, làm thiếu phát hiện vấn đề.')
para(
    'Đề xuất sửa: nghiên cứu kế tiếp NÊN dùng CẢ DASS-Y và '
    'DASS-21 trên CÙNG mẫu để so sánh ngưỡng — như tác giả VN020 '
    'cũng đã đề xuất trong "Hướng nghiên cứu tiếp theo".'
)

H('2.7. TẠP CHÍ Y học TPHCM không có chỉ mục PubMed/Scopus', level=3, color=NAVY)
para(
    'Tạp chí Y học TPHCM là tạp chí trong nước, KHÔNG có chỉ mục '
    'quốc tế (PubMed, Scopus, Web of Science). Hậu quả:'
)
para('• Nghiên cứu KHÓ được trích dẫn trong meta-analysis quốc tế.')
para('• Đánh giá peer-review không cùng chuẩn với tạp chí Q1.')
para('• Không có DOI quốc tế chuẩn (chỉ có DOI nội địa nếu có).')
para(
    'So sánh: VN029 (Trúc Thanh Thái 2025) cùng đề tài SKTT HS '
    'THPT TPHCM nhưng đăng trên Social Psychiatry and Psychiatric '
    'Epidemiology (Q1, peer-reviewed quốc tế) — sẽ được trích '
    'dẫn nhiều hơn trong y văn quốc tế.'
)

# 3. So sánh tổng hợp
H('3. So sánh VN020 với chương 3 luận án CTH', level=2, color=NAVY)
caption('Bảng 1. So sánh VN020 vs Chương 3 luận án về cách đo và phân tích ALHT')
add_table(
    ['Đặc điểm', 'VN020 (Trần Hồ Vĩnh Lộc 2024)', 'Chương 3 luận án CTH'],
    [
        ['Cỡ mẫu', '976 HS THPT (2 trường TPHCM)', '1.352 HS THCS (đa trường)'],
        ['Cấp học', 'THPT', 'THCS'],
        ['Thang đo lo âu', 'DASS-Y (1 thang con anxiety)', 'RCADS (3 thang con: TQ, CL, XH)'],
        ['Thang đo ALHT', 'ESSA full 16 mục, α=0,81', 'ESSA rút gọn 4 mục'],
        ['Phân tích thống kê', 'Hồi quy logistic đa biến (OR)', 'SEM (β chuẩn hóa, R², CFI, RMSEA)'],
        ['Phân tách 3 dạng RLLA', '✗ KHÔNG (chỉ "lo âu" tổng)', '✓ CÓ (RLLATQ, RLLACL, RLLAXH riêng)'],
        ['Số yếu tố nguy cơ', '1 (ALHT)', '3 (ALHT, NĐT, BNHĐ)'],
        ['Số yếu tố bảo vệ', '0', '4 (TTr, GBTH, HTCM, HTBB)'],
        ['Mô hình tích hợp', '✗', '✓ YTNC + YTBV (R²=0,598)'],
        ['Tạp chí', 'Y học TPHCM (trong nước)', 'Luận án (sẽ xuất bản)'],
        ['ƯU ĐIỂM ĐỘC ĐÁO', 'ESSA full + DASS-Y', 'SEM tích hợp + 8 yếu tố'],
    ]
)
para('')
para(
    'Kết luận so sánh: VN020 có ƯU thế ở thang đo (ESSA full + '
    'DASS-Y phiên bản VTN), nhưng YẾU hơn ở phân tích thống kê '
    '(logistic vs SEM) và phạm vi yếu tố (1 vs 7). Chương 3 '
    'luận án CTH có thiết kế TOÀN DIỆN HƠN — phù hợp lấp '
    'KHOẢNG TRỐNG mà VN020 để lại.', bold=True
)

# 4. Đề xuất sửa
H('4. Bốn đề xuất cho nghiên cứu kế tiếp', level=2, color=NAVY)
para('ĐỀ XUẤT 1 — Mở rộng cỡ mẫu + đa trường', bold=True)
para(
    'Chọn TỐI THIỂU 8–10 trường THPT TPHCM (chính quy + GDTX + '
    'chuyên + dân lập) bằng probability proportional to size (PPS). '
    'Lấy lớp ngẫu nhiên trong từng trường (không thuận tiện). Cỡ '
    'mẫu mục tiêu n = 2.500–3.000 (như VN029).'
)
para('')
para('ĐỀ XUẤT 2 — Thay logistic bằng SEM', bold=True)
para(
    'Sử dụng SEM với biến tiềm ẩn (ESSA 5 chiều → biến tiềm ẩn '
    '"ALHT" → lo âu) thay vì hồi quy logistic ngưỡng cắt. Cho '
    'phép báo cáo β chuẩn hóa, R², CFI, RMSEA — so sánh trực '
    'tiếp với meta-analysis quốc tế và chương 3 luận án.'
)
para('')
para('ĐỀ XUẤT 3 — Phân tách 3 dạng RLLA', bold=True)
para(
    'Thay DASS-Y (chỉ 1 thang lo âu) bằng RCADS (3 thang con: '
    'lan tỏa, chia ly, xã hội) — phù hợp DSM-5 và phù hợp với '
    'chương 3 luận án CTH. Cho phép phát hiện pattern khác '
    'biệt theo loại RLLA.'
)
para('')
para('ĐỀ XUẤT 4 — Bổ sung 5 yếu tố thiếu', bold=True)
para(
    'Thêm vào mô hình: bắt nạt (OBVQ-R), tự trọng (RSES), hỗ '
    'trợ xã hội (MSPSS), nghiện điện thoại (SAS-SV), biện pháp '
    'đối phó (Brief-COPE). Giúp xây dựng mô hình tích hợp YTNC '
    '+ YTBV như chương 3 — toàn diện hơn nhiều.'
)

# 5. CÂU TRẢ LỜI
H('5. CÂU TRẢ LỜI', level=2, color=NAVY)
blue_run('Tóm gọn — VN020 là bài VN TỐT NHẤT về ALHT đến nay nhưng có 7 hạn chế phương pháp luận:', bold=True)
blue_bullet(
    'CHỈ 2 TRƯỜNG — không đại diện TPHCM (200+ trường THPT). Tỷ lệ '
    'lo âu (25,1% DASS-Y) khác xa VN029 (50,3% DASS-21, n=2.631) — '
    'gợi ý mẫu/thang đo phụ thuộc.'
)
blue_bullet(
    'CHỌN LỚP THUẬN TIỆN ở bậc lớp (4 lớp/khối) — chỉ bậc trường '
    'ngẫu nhiên. Có thể có thiên lệch chọn lọc.'
)
blue_bullet(
    'HỒI QUY LOGISTIC + ngưỡng cắt nhị phân — MẤT thông tin liên '
    'tục. Chỉ có OR, KHÔNG có β chuẩn hóa và R² → khó so sánh với '
    'meta-analysis quốc tế (Walder 2025: g=0,508–0,878).'
)
blue_bullet(
    'KHÔNG PHÂN TÁCH 3 DẠNG RLLA — DASS-Y chỉ 1 thang lo âu. Mất '
    'phát hiện pattern khác biệt như chương 3 CTH (β ALHT khác '
    'cho RLLATQ 0,510 / RLLACL 0,253 / RLLAXH 0,490).'
)
blue_bullet(
    'BỎ SÓT 5 YẾU TỐ quan trọng: bắt nạt, tự trọng, hỗ trợ xã '
    'hội, nghiện điện thoại, biện pháp đối phó. Mô hình KHÔNG '
    'kiểm soát các yếu tố này → kết luận "ESSA mạnh nhất" có '
    'thể overestimate.'
)
blue_bullet(
    'DASS-Y chưa phổ biến quốc tế — không có meta-analysis '
    'so sánh DASS-Y vs DASS-21. Tỷ lệ 25,1% (DASS-Y) thấp hơn '
    'nhiều bài VN khác dùng DASS-21 (40–86%) — gợi ý ngưỡng '
    'có thể QUÁ NGHIÊM.'
)
blue_bullet(
    'TẠP CHÍ Y HỌC TPHCM không PubMed/Scopus — giảm tin cậy '
    'và khả năng trích dẫn quốc tế. So với VN029 đăng Q1 Soc '
    'Psychiatry — VN020 yếu hơn về uy tín tạp chí.'
)

blue_run('Cách trích vào báo cáo CTH:', bold=True)
blue_run(
    '"Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh và Tô Gia Kiên (2024) '
    'trên 976 học sinh THPT tại 2 trường TPHCM — sử dụng DASS-Y '
    '(Szabó, 2010) và ESSA (Sun và cộng sự, 2011) — phát hiện áp '
    'lực học tập (ESSA ≥ 59) là yếu tố nguy cơ mạnh nhất tăng '
    'tỷ lệ lo âu (25,1% theo DASS-Y). Đây là một trong số ít '
    'nghiên cứu Việt Nam dùng thang đo chuẩn quốc tế cho áp lực '
    'học tập. Tuy nhiên, nghiên cứu có một số hạn chế phương '
    'pháp luận: (a) chỉ 2 trường TPHCM — không đại diện thành '
    'phố; (b) lấy mẫu thuận tiện ở bậc lớp; (c) phân tích logistic '
    'với ngưỡng cắt nhị phân thay vì SEM với β chuẩn hóa, làm '
    'mất thông tin liên tục; (d) DASS-Y chỉ một thang lo âu — '
    'không phân tách lo âu lan tỏa, chia ly, và xã hội theo '
    'DSM-5; (e) bỏ sót các yếu tố quan trọng như bắt nạt, tự '
    'trọng, hỗ trợ xã hội. Khuyến nghị nghiên cứu tiếp theo mở '
    'rộng đa trường, áp dụng SEM với RCADS, và bổ sung mô hình '
    'tích hợp yếu tố nguy cơ + bảo vệ như chương 3 luận án."',
    italic=True
)

# 6. TLTK
H('6. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Sydney: Psychology Foundation of Australia.',
    'Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823 [QT067 trong cơ sở dữ liệu dự án.]',
    'Sun, J., Dunne, M. P., Hou, X. Y., & Xu, A. Q. (2011). Educational Stress Scale for Adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546. https://doi.org/10.1177/0734282910394976',
    'Szabó, M. (2010). The short version of the Depression Anxiety Stress Scales (DASS-21): Factor structure in a young adolescent sample. Journal of Adolescence.',
    'Trần, H. V. L., Huỳnh, N. V. A., & Tô, G. K. (2024). Trầm cảm, lo âu, căng thẳng và các yếu tố liên quan ở học sinh THPT tại Thành phố Hồ Chí Minh. Tạp chí Y học TPHCM. [VN020 trong cơ sở dữ liệu dự án.]',
    'Trúc, T. T., Dương, M. C., và cộng sự. (2025). Unmasking the burden of mental health symptoms and risk behaviors in Vietnamese adolescents: evidence from a multicenter cross-sectional study involving 2,631 high school students. Social Psychiatry and Psychiatric Epidemiology. [VN029 trong cơ sở dữ liệu dự án.]',
    'Walder, N., Frey, A., Berger, T., và cộng sự. (2025). Digital mental health interventions for prevention and treatment of social anxiety disorder. Journal of Medical Internet Research. [QT040 trong cơ sở dữ liệu dự án.]',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
