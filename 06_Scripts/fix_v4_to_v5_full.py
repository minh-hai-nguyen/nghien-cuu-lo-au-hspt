# -*- coding: utf-8 -*-
"""
FULL REWORK v4 -> v5: sửa fabrication trong 2 bài VJES.
- Rebuild TLTK VN entries với dữ liệu ĐÃ VERIFY từ PDF gốc (2 agent đọc PDF 15/05/2026)
  + VN005 verify từ RAG/Tóm tắt (PDF là bản scan ảnh, không trích text được).
- Viết lại các đoạn thân bài gán phát hiện sai cho bài VN.
- Chỉ trích dẫn ĐÚNG điều mỗi bài VN thật sự phát hiện + đúng bối cảnh (THPT/THCS).
- Chuẩn hóa "& cs" -> "và cộng sự" (VJES: không viết tắt).
v4 được GIỮ NGUYÊN để đối chứng — script chỉ đọc v4, ghi ra file v5 riêng.
"""
import os
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn"
B1_IN = os.path.join(ROOT, "Bai1_YTNC_HSTHCS_v4_14052026.docx")
B2_IN = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v4_14052026.docx")
B1_OUT = os.path.join(ROOT, "Bai1_YTNC_HSTHCS_v5_15052026.docx")
B2_OUT = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v5_15052026.docx")


def replace_para_keep_format(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def replace_by_unique(doc, unique_substr, new_text):
    """Tìm đoạn chứa unique_substr, thay TOÀN BỘ text đoạn bằng new_text."""
    for p in doc.paragraphs:
        if unique_substr in p.text:
            replace_para_keep_format(p, new_text)
            return True
    return False


def fix_abbrev(doc):
    """Chuẩn hóa '& cs' -> 'và cộng sự' trong toàn bộ đoạn văn."""
    n = 0
    for p in doc.paragraphs:
        if "& cs" in p.text:
            for r in p.runs:
                if "& cs" in r.text:
                    r.text = r.text.replace("& cs", "và cộng sự")
                    n += 1
            # phòng trường hợp text nằm rải nhiều run
            if "& cs" in p.text:
                replace_para_keep_format(p, p.text.replace("& cs", "và cộng sự"))
                n += 1
    return n


# ============================================================
# BÀI 1 — body paragraph rewrites
# ============================================================
B1_BODY = [
    # §1 — Trúc Thanh Thái: năm 2026 (PDF: © 2026, advance online)
    ("mẫu 2.631 học sinh trung học TP. Hồ Chí Minh của Dương",
     "Tại Việt Nam, Điều tra Sức khỏe tâm thần vị thành niên quốc gia (V-NAMHS) do Viện Xã hội học phối hợp Đại học Queensland và Đại học Johns Hopkins công bố năm 2022 trên 5.996 cặp phụ huynh – vị thành niên 10–17 tuổi cho thấy khoảng 21,7% mẫu có ít nhất một vấn đề sức khỏe tâm thần trong 12 tháng qua, trong khi tỷ lệ rối loạn theo tiêu chí chẩn đoán DSM-5 đầy đủ cho nhóm lo âu là 2,3% (UNICEF Việt Nam, 2022). Khoảng cách giữa tỷ lệ vấn đề và tỷ lệ chẩn đoán phản ánh sự khác biệt giữa sàng lọc và chẩn đoán lâm sàng. Khảo sát 8.473 học sinh khối 6–12 tại sáu tỉnh của Hoàng Trung Học và Nguyễn Thùy Dung (2025) bằng thang Đánh giá Trầm cảm – Lo âu – Stress phiên bản 21 mục (DASS-21) ghi nhận tỷ lệ triệu chứng lo âu ở mức sàng lọc là 41,5% trong giai đoạn đại dịch COVID-19 và giảm còn 25,4% sau đại dịch — cao hơn nhiều lần tỷ lệ chẩn đoán DSM-5. Khảo sát 2.631 học sinh trung học phổ thông Thành phố Hồ Chí Minh của Trúc Thanh Thái và cộng sự (2026) cũng ghi nhận tỷ lệ triệu chứng lo âu cao, ở mức 50,3%."),

    # §3.1 — bỏ claim giới (nguồn tự mâu thuẫn); Mỵ Lương; số liệu đã verify
    ("công trình của Đinh và cộng sự (2021) trên hơn 1.300 học sinh THPT",
     "Tại Việt Nam, công trình của Đinh Thị Hồng Vân và cộng sự (2021) trên 749 học sinh trung học cơ sở tại Hà Nội và Thành phố Hồ Chí Minh cho thấy điểm số và thi cử là yếu tố gây lo âu được ghi nhận ở 91,1% học sinh — cao nhất trong sáu nhóm yếu tố trường học được khảo sát. Trần Thị Mỵ Lương (2020) trên 540 học sinh trung học phổ thông chuyên (thang DASS-42) ghi nhận tỷ lệ lo âu 14,2%, trong đó học sinh khối 11 chiếm tỷ lệ cao nhất (48,1% trong số ca lo âu), được lý giải bằng áp lực đa chiều của giai đoạn chuyển tiếp này. Khảo sát 3.910 học sinh trung học phổ thông Hà Nội của Phạm Thị Thu Hoa và cộng sự (2024) bằng thang Rối loạn Lo âu Tổng quát 7 mục (GAD-7) ghi nhận 40,6% học sinh có nguy cơ lo âu, trong đó học tập được xác định là nguyên nhân hàng đầu. Mặc dù phần lớn dữ liệu Việt Nam thuộc bậc trung học phổ thông, hướng của hiệu ứng nhất quán với bằng chứng quốc tế và gợi ý áp lực học tập là yếu tố nguy cơ ổn định xuyên cấp học."),

    # §3.2 — bỏ claim bịa Hoa-điện thoại + Phạm Thị Ngọc-SAS; dùng VN014 thật
    ("công trình của Hoa và cộng sự (2024) ghi nhận thời gian sử dụng điện thoại trên ba giờ",
     "Tại Việt Nam, khảo sát 8.473 học sinh của Hoàng Trung Học và Nguyễn Thùy Dung (2025) xác định việc sử dụng thiết bị điện tử kéo dài là một trong các yếu tố liên hệ thuận với mức độ lo âu, được ghi nhận cả trong và sau đại dịch COVID-19. Tuy nhiên, các nghiên cứu Việt Nam đo lường trực tiếp nghiện điện thoại thông minh bằng công cụ chuẩn hóa trên học sinh trung học cơ sở vẫn còn rất hạn chế — đây là một khoảng trống cần được lấp đầy. Bằng chứng quốc tế hiện có đã đủ để xếp nghiện điện thoại thông minh vào nhóm yếu tố nguy cơ trọng yếu cần được ưu tiên khảo sát trong bối cảnh Việt Nam."),

    # §3.3 — VN025 (bạo lực aOR=2,42) + VN015 (quan hệ bạn bè + ACE) thật
    ("báo cáo của UNICEF Việt Nam (2022) khảo sát tại Gia Lai và TP. Hồ Chí Minh cho thấy tỷ lệ học sinh từng là nạn nhân bắt nạt thể chất",
     "Tại Việt Nam, khảo sát 420 học sinh trung học phổ thông và giáo dục thường xuyên huyện Vĩnh Bảo, Hải Phòng của Phạm Thị Ngọc và cộng sự (2024) ghi nhận học sinh từng bị bạo lực học đường có nguy cơ lo âu cao gấp 2,42 lần so với nhóm không bị (tỷ số chênh hiệu chỉnh; khoảng tin cậy 95% từ 1,03 đến 5,71). Nghiên cứu 845 học sinh dân tộc thiểu số nội trú tỉnh Lạng Sơn của Ngô Anh Vinh và cộng sự (2024) cũng cho thấy chất lượng quan hệ bạn bè kém có liên hệ thuận với cả lo âu, stress và trầm cảm, đồng thời số trải nghiệm bất lợi thời thơ ấu liên hệ thuận với điểm lo âu. Báo cáo của UNICEF Việt Nam (2022) ghi nhận bắt nạt học đường là vấn đề phổ biến ở học sinh Việt Nam. Hướng của các phát hiện trùng với bằng chứng quốc tế, đồng thời gợi ý rằng tại Việt Nam, bắt nạt có thể tương tác với khoảng cách kinh tế – xã hội theo vùng miền — một biến mà các nghiên cứu hiện có chưa khai thác đầy đủ."),

    # §3.4 — bỏ claim bịa Hoàng Trung Học verbal; dùng VN025 (không hòa đồng aOR=2,25)
    ("Hoàng Trung Học và cộng sự (2025) xác nhận trải nghiệm bị bắt nạt bằng lời",
     "Tại Việt Nam, khảo sát của Phạm Thị Ngọc và cộng sự (2024) ghi nhận học sinh không có quan hệ hòa đồng với bạn bè có nguy cơ lo âu cao gấp 2,25 lần so với nhóm hòa đồng — phản ánh tác động của loại trừ xã hội, một dạng tổn thương gần với bắt nạt bằng lời nói. Các hình thức bắt nạt bằng lời thường được phối hợp với hành vi loại trừ xã hội — chẳng hạn cô lập trong nhóm bạn lớp, từ chối ngồi gần hoặc lập nhóm trò chuyện riêng có chủ ý loại trừ một học sinh. Khi cộng dồn các hành vi này, tổn thương tâm lý có thể vượt xa tổn thương từ bắt nạt thể chất đơn lẻ, ngay cả khi không để lại bằng chứng vật lý. Bắt nạt mạng làm gia tăng mức độ phơi nhiễm do tính ẩn danh và khả năng lan tỏa nhanh trong môi trường trực tuyến."),

    # §3.5 — VN003 (sửa: 9 tỉnh không phải Huế; bỏ claim chăm sóc thể chất) + VN020
    ("Phạm và cộng sự (2024) trên HSTHCS Huế bằng MSPSS và thang đo lo âu chuyên biệt",
     "Tại Việt Nam, bằng chứng định lượng trực tiếp về lòng tự trọng và lo âu ở học sinh trung học cơ sở còn rất hạn chế. Nghiên cứu của Phạm Sỹ Tiến và cộng sự (2024) trên 546 thanh thiếu niên 12–18 tuổi tại các cơ sở hỗ trợ xã hội thuộc chín tỉnh, thành phố cho thấy chất lượng chăm sóc cảm xúc có liên hệ nghịch và đáng kể với các vấn đề sức khỏe tâm thần, trong đó có lo âu (hệ số β = −0,40; p < 0,001) — gợi ý vai trò của nâng đỡ cảm xúc và tự đánh giá tích cực đối với sức khỏe tâm thần của thanh thiếu niên. Khảo sát 976 học sinh trung học phổ thông Thành phố Hồ Chí Minh của Trần Hồ Vĩnh Lộc và cộng sự (2024) bằng thang DASS-Y xác định áp lực học tập là yếu tố nguy cơ mạnh đối với lo âu, với tỷ số tỷ lệ hiện mắc hiệu chỉnh đạt 2,82 ở nhóm chịu áp lực học tập mức nặng. Khoảng trống về đo lường trực tiếp lòng tự trọng trong các nghiên cứu trên học sinh trung học cơ sở Việt Nam cần được ưu tiên lấp đầy."),
]


# ============================================================
# BÀI 1 — TLTK rebuild (10 mục VN, đã verify từ PDF 15/05/2026)
# ============================================================
B1_TLTK = [
    # VN029 — năm 2026 (PDF: © The Author(s) 2026, advance online)
    ("Duong, T. T. T., Tran, M. T., & Nguyen, H. P. (2025)",
     "Trúc Thanh Thái, Võ Lê Hồng Tuyết, Nguyễn Thị Trang, Đinh Văn Ngôn, Mai Lê Xuân, Trần Thị Hoài Thương, Nguyễn Thị Ngọc Bích, Huỳnh Mai Khánh Hà, Nguyễn Thị Thu An, Bùi Thị Hy Hân, & Dương Minh Cường. (2026). Unmasking the burden of mental health symptoms and risk behaviors in Vietnamese adolescents: Evidence from a multicenter cross-sectional study involving 2,631 high school students. Social Psychiatry and Psychiatric Epidemiology. Xuất bản trực tuyến trước. https://doi.org/10.1007/s00127-025-03043-7"),

    # VN027
    ("Đinh Thị Hồng Vân. (2021)",
     "Đinh Thị Hồng Vân, Đỗ Thị Lê Hằng, & Phan Thị Mai Hương. (2021). School factors causing Vietnamese adolescents' anxiety in secondary schools. Psychology and Education, 58(1), 883–894."),

    # VN001
    ("Hoa, T. T. H., Phan, P. H., & Tran, T. K. (2024)",
     "Phạm Thị Thu Hoa, Đỗ Thị Trang, Nguyễn Thị Liên, & Ngô Anh Vinh. (2024). Anxiety symptoms and coping strategies among high school students in Vietnam after COVID-19 pandemic: A mixed-method evaluation. Frontiers in Public Health, 12, Article 1232856. https://doi.org/10.3389/fpubh.2024.1232856"),

    # VN014
    ("Hoàng Trung Học, Nguyễn Thanh Bình, & Trần Thị Thu Hằng. (2025)",
     "Hoàng Trung Học, & Nguyễn Thùy Dung. (2025). Levels of stress, anxiety, and depression in adolescents during and after the COVID-19 pandemic in Vietnam: A cross-sectional study. American Journal of Psychiatric Rehabilitation, 28(1), 360–367. https://doi.org/10.69980/ajpr.v28i1.105"),

    # VN015
    ("Ngô Anh Vinh, Phạm Mạnh Hùng, & Đỗ Thị Hồng Liên. (2024)",
     "Ngô Anh Vinh, Vũ Thị Mỹ Hạnh, Đỗ Thị Bích Vân, Dương Anh Tài, Đỗ Minh Loan, & Lê Thị Thanh Thùy. (2024). Mental health among ethnic minority adolescents in Vietnam and correlated factors: A cross-sectional study. Journal of Affective Disorders Reports, 17, Article 100795. https://doi.org/10.1016/j.jadr.2024.100795"),

    # VN003 — sửa tên tạp chí (bỏ "the")
    ("Phạm, T. M. H., Nguyễn, T. K. L., & Lê, V. A. (2024)",
     "Phạm Sỹ Tiến, Dương Thị Thanh Thanh, Nguyễn Thị Hoài Phương, & Trương Thị Xuân Nhi. (2024). The correlation between quality of care and mental health and behavioral problems among Vietnamese adolescents in social support facilities. Journal of Indian Association for Child and Adolescent Mental Health, 20(4), 373–381. https://doi.org/10.1177/09731342241275742"),

    # VN025 — bổ sung trang 115-123
    ("Phạm Thị Ngọc, Lê Quang Sơn, & Trần Văn Hùng. (2024)",
     "Phạm Thị Ngọc, Hoàng Thị Giang, Phạm Khánh Linh, & Vũ Thị Châu. (2024). Thực trạng sức khỏe tâm thần và một số yếu tố liên quan của học sinh tại 2 trường trung học phổ thông huyện Vĩnh Bảo, Hải Phòng năm 2023. Tạp chí Y học dự phòng, 34(1), 115–123. https://doi.org/10.51403/0868-2836/2024/1571"),

    # VN020
    ("Trần Hồ Vĩnh Lộc. (2024)",
     "Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh, & Tô Gia Kiên. (2024). Trầm cảm, lo âu, căng thẳng và các yếu tố liên quan ở học sinh trung học phổ thông tại Thành phố Hồ Chí Minh. Tạp chí Y học Thành phố Hồ Chí Minh, 27(5), 100–110. https://doi.org/10.32895/hcjm.m.2024.05.12"),

    # VN021
    ("Trần Thảo Vi, Nguyễn Văn Tuấn, & Hoàng Thị Loan. (2024)",
     "Trần Thảo Vi, Nguyễn Hoàng Thùy Linh, Trần Xuân Minh Trí, Tashiro, Y., Seino, K., Võ Văn Thắng, & Nakamura, K. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine, 19(4), 279–290. https://doi.org/10.2185/jrm.2024-012"),

    # VN006 — sửa "Mỵ" (không phải "Mỹ"), tên tạp chí, trang 122-131
    ("Trần Thị Mỵ Lương, & Phạm Thị Thu Hương. (2020)",
     "Trần Thị Mỵ Lương. (2020). Rối loạn lo âu ở học sinh trung học phổ thông. Tạp chí Khoa học, Trường Đại học Thủ đô Hà Nội, 40, 122–131."),
]


# ============================================================
# BÀI 2 — body + TLTK
# ============================================================
B2_BODY = [
    # §3.1 — VN005 (Trần Nguyễn Ngọc, RLLA lan tỏa) + Happy House (Tran và cộng sự 2023)
    ("Luận án tiến sĩ của Trần Nguyễn Ngọc (2018) tại Bệnh viện Bạch Mai",
     "Tại Việt Nam, các nghiên cứu can thiệp dựa trên CBT trên học sinh trung học cơ sở còn rất hạn chế. Luận án tiến sĩ y học của Trần Nguyễn Ngọc (2018) tại Trường Đại học Y Hà Nội đánh giá hiệu quả của liệu pháp thư giãn luyện tập — một thành phần hành vi của CBT — trên bệnh nhân rối loạn lo âu lan tỏa; mặc dù mẫu là người trưởng thành chứ không phải học sinh trung học cơ sở, công trình cung cấp dữ liệu về hiệu quả của kỹ thuật thư giãn trong bối cảnh Việt Nam. Chương trình Happy House do Tran và cộng sự (2023) thử nghiệm trong khuôn khổ hợp tác Việt – Úc đã đánh giá một can thiệp nâng cao sức khỏe tâm thần phổ quát theo mô hình thử nghiệm hai nhánh có đối chứng trên 1.084 học sinh trung học phổ thông tại Hà Nội. Tuy nhiên, các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu học sinh trung học cơ sở Việt Nam vẫn chưa được triển khai."),
]

B2_TLTK = [
    # VN005 — sửa tên "Nguyễn" (không phải "Nguyên"), thêm "lan tỏa", nguồn
    ("Trần Nguyễn Ngọc. (2018)",
     "Trần Nguyễn Ngọc. (2018). Đánh giá hiệu quả điều trị rối loạn lo âu lan tỏa bằng liệu pháp thư giãn luyện tập [Luận án tiến sĩ y học, Trường Đại học Y Hà Nội]. Trường Đại học Y Hà Nội."),

    # Happy House — sửa DOI/volume/article (verify web 15/05: vol 10, e69, gmh.2023.66)
    ("Trần, T. M. T., Doan, V. H., & Pham, A. T. (2023)",
     "Tran, T. D., Nguyen, H., Shochet, I., Nguyen, N., La, N., Wurfl, A., Orr, J., Nguyen, H., Stocker, R., & Fisher, J. (2023). School-based universal mental health promotion intervention for adolescents in Vietnam: Two-arm, parallel, controlled trial. Cambridge Prisms: Global Mental Health, 10, Article e69. https://doi.org/10.1017/gmh.2023.66"),
]


def process(in_path, out_path, body_fixes, tltk_fixes, label):
    print(f"=== {label} ===")
    doc = Document(in_path)
    nb = 0
    for sub, new in body_fixes:
        if replace_by_unique(doc, sub, new):
            nb += 1
            print(f"  OK body: ...{sub[:50]}...")
        else:
            print(f"  !! body NOT FOUND: ...{sub[:50]}...")
    nt = 0
    for sub, new in tltk_fixes:
        if replace_by_unique(doc, sub, new):
            nt += 1
            print(f"  OK TLTK: ...{sub[:45]}...")
        else:
            print(f"  !! TLTK NOT FOUND: ...{sub[:45]}...")
    na = fix_abbrev(doc)
    print(f"  -> body {nb}/{len(body_fixes)}, TLTK {nt}/{len(tltk_fixes)}, '& cs' fixed: {na}")

    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.subject = ""
    cp.category = ""
    cp.title = ""
    cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 15, 10, 0, 0)
    doc.save(out_path)
    print(f"  Saved: {out_path}\n")


if __name__ == "__main__":
    process(B1_IN, B1_OUT, B1_BODY, B1_TLTK, "BAI 1")
    process(B2_IN, B2_OUT, B2_BODY, B2_TLTK, "BAI 2")
    print("[DONE]")
