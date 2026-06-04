"""Build doc tra loi: OR=11,579 KTC 95% [4,164; 32,194] - khoang co du manh khong?"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/OR_11579_KTC_95_4_164_32_194_do_manh.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(12); r.italic = italic
    return p

# =====================================================
# TIÊU ĐỀ
# =====================================================
H('Bình luận về độ mạnh và độ chính xác của OR = 11,579', level=1)
H('— KTC 95%: 4,164; 32,194 — khoảng có đủ mạnh không? —', level=2)

# =====================================================
# CÂU HỎI (xanh)
# =====================================================
H('Câu hỏi của thầy', level=2)
para_blue(
    'Cái này về khoảng có đủ mạnh không? (OR = 11,579; KTC 95% [4,164; 32,194]).'
)
para_black('')
para_black(
    'Bối cảnh: số liệu trích từ Wen et al. (2020) "Latent profile analysis of anxiety '
    'among junior high school students in less developed rural areas of China" '
    '(QT008 trong DB dự án), hệ số OR cho áp lực học tập rất cao ↔ lo âu nặng '
    'so với áp lực rất thấp.', italic=True
)

# =====================================================
# BACKGROUND (đen)
# =====================================================
H('1. Phân biệt 3 khái niệm: TÍNH Ý NGHĨA, ĐỘ MẠNH, ĐỘ CHÍNH XÁC', level=2)

para_black('Khi đánh giá một KTC (khoảng tin cậy), người đọc thường nhầm 3 khái niệm khác nhau:')

tbl = doc.add_table(rows=4, cols=3)
tbl.style = 'Light Grid Accent 1'
header = tbl.rows[0]
header.cells[0].text = 'Khái niệm'
header.cells[1].text = 'Đo bằng'
header.cells[2].text = 'Áp dụng vào OR=11,579 [4,16; 32,19]'
data = [
    ('Tính ý nghĩa thống kê', 'KTC có CHỨA giá trị null (1 cho OR/RR) không?',
     'KTC [4,16; 32,19] KHÔNG chứa 1 → CÓ ý nghĩa thống kê (p << 0,05).'),
    ('Độ MẠNH (effect size)', 'Mức độ lệch khỏi null. Cận DƯỚI là quan trọng nhất.',
     'Cận dưới 4,16 → ngay cả "tệ nhất", OR vẫn ≥ 4 → MẠNH (>4 = strong; >6 = very strong).'),
    ('Độ CHÍNH XÁC (precision)', 'Width của KTC: cận trên − cận dưới.',
     'Width = 32,19 − 4,16 = 28,03 → RẤT RỘNG → độ chính xác THẤP.'),
]
for i, (a, b, c) in enumerate(data, 1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black(
    'Một OR = 11,5 với KTC 95% [4,16; 32,19] cho biết: chắc chắn có liên quan dương, '
    'mức ÍT NHẤT là 4 lần, nhưng tác giả KHÔNG biết chắc OR thực là 5, 11, hay 30 — '
    'đây là khoảng RẤT rộng, biểu thị mẫu KHÔNG đủ lớn để ước lượng chính xác.', italic=True
)

H('2. Tại sao KTC lại rộng đến vậy?', level=2)

para_black('KTC rộng gặp khi 1 hoặc nhiều yếu tố sau xảy ra:')
bullet_black(
    'Cỡ mẫu nhỏ (n < 1000) — sai số chuẩn (SE) lớn nên KTC nở ra. Wen 2020 có n=900 '
    'với 6 mức "áp lực học tập" (Likert) — sau khi chia subgroup, một số ô có thể chỉ '
    'còn 30–50 quan sát.'
)
bullet_black(
    'Outcome HIẾM ở nhóm tham chiếu (reference group). Ở đây "lo âu nặng" trong nhóm '
    '"áp lực rất thấp" thường chỉ vài %. Nếu chỉ 1–2 ca trong ô tham chiếu, OR sẽ phình to.'
)
bullet_black(
    'Outcome PHỔ BIẾN ở nhóm phơi nhiễm. Nếu "áp lực rất cao" + "lo âu nặng" ~ 30–50%, '
    'odds nhóm này lớn → tỷ số OR phình.'
)
bullet_black(
    '"Zero-cell problem" — khi 1 trong 4 ô của bảng 2×2 = 0, các phần mềm thường thêm '
    '0,5 vào tất cả ô (Haldane-Anscombe correction) → OR ước lượng méo + KTC rộng cực đoan.'
)

H('3. Phân loại "rộng" của KTC OR', level=2)

tbl2 = doc.add_table(rows=5, cols=3)
tbl2.style = 'Light Grid Accent 1'
header = tbl2.rows[0]
header.cells[0].text = 'Tỷ số cận trên/cận dưới (UL/LL)'
header.cells[1].text = 'Phân loại'
header.cells[2].text = 'OR = 11,58 [4,16; 32,19]'
rows = [
    ('< 2', 'KTC HẸP — độ chính xác cao', ''),
    ('2 – 4', 'KTC trung bình', ''),
    ('4 – 8', 'KTC RỘNG — độ chính xác thấp', '✓ UL/LL = 32,19/4,16 = 7,74 → mức RỘNG'),
    ('> 8', 'KTC RẤT RỘNG — chỉ phù hợp đặt giả thuyết', ''),
]
for i, (a, b, c) in enumerate(rows, 1):
    tbl2.rows[i].cells[0].text = a
    tbl2.rows[i].cells[1].text = b
    tbl2.rows[i].cells[2].text = c
for row in tbl2.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black(
    '⚠ LƯU Ý: bảng phân loại UL/LL trên là QUY ƯỚC THỰC HÀNH em tổng hợp từ kinh nghiệm '
    'đọc các paper dịch tễ — KHÔNG phải chuẩn cố định từ một paper duy nhất. Nguyên tắc '
    'cốt lõi (Altman & Bland 1995, BMJ; Greenland et al. 2016, Eur J Epidemiol): KTC '
    'rộng = ước lượng kém chính xác; người đọc cần nhìn cận DƯỚI để biết "tệ nhất" '
    'effect size là bao nhiêu, không nên dùng riêng ước lượng điểm.', italic=True
)

H('4. So sánh với các OR khác trong corpus dự án', level=2)

tbl3 = doc.add_table(rows=8, cols=4)
tbl3.style = 'Light Grid Accent 1'
header = tbl3.rows[0]
header.cells[0].text = 'Bài'
header.cells[1].text = 'OR'
header.cells[2].text = 'KTC 95%'
header.cells[3].text = 'Diễn giải'
rows3 = [
    ('Wen 2020 (QT008) — áp lực rất cao ↔ LA nặng', '11,58', '4,16; 32,19',
     'Mạnh (LL>4) nhưng KTC rất rộng → mẫu nhỏ'),
    ('Qiu 2022 (QT008) — phục hồi thấp ↔ trầm cảm', '6,74', '4,73; 9,61',
     'Rất mạnh + KTC HẸP → đáng tin cậy'),
    ('Qiu 2022 — phục hồi thấp ↔ lo âu', '2,80', '1,92; 4,09',
     'Trung bình + KTC trung bình'),
    ('Qiu 2022 — nuôi dạy tiêu cực ↔ lo âu', '2,01', '1,38; 2,92',
     'Yếu-TB + KTC trung bình'),
    ('Qiu 2022 — nuôi dạy tích cực ↔ lo âu', '0,32', '0,24; 0,43',
     'Bảo vệ TB + KTC HẸP → đáng tin'),
    ('Zhu 2025 (QT trong corpus) — thiếu ngủ ↔ trầm cảm', '13,71', 'N/A',
     'Rất mạnh — cần check KTC'),
    ('Anderson 2025 (QT014) — áp lực ↔ lo âu', 'N/A', 'N/A',
     'Narrative review — không pooled OR'),
]
for i, (a, b, c, d) in enumerate(rows3, 1):
    tbl3.rows[i].cells[0].text = a
    tbl3.rows[i].cells[1].text = b
    tbl3.rows[i].cells[2].text = c
    tbl3.rows[i].cells[3].text = d
for row in tbl3.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10); r.font.name = 'Times New Roman'

para_black('')
para_black(
    'Wen 2020 OR=11,58 (KTC width 28) so với Qiu 2022 OR=6,74 (KTC width 4,88): '
    'KTC của Wen rộng GẤP 5,7 LẦN của Qiu — Qiu đáng tin cậy hơn về precision dù effect '
    'size nhỏ hơn.', italic=True
)

H('5. Cẩn trọng khi áp dụng OR=11,58 vào báo cáo', level=2)
bullet_black(
    'Cận dưới 4,16 vẫn là MẠNH → vẫn có thể trích để khẳng định "có liên quan rất mạnh".'
)
bullet_black(
    'KHÔNG trích "OR = 11,58" như con số tuyệt đối; phải kèm KTC để người đọc thấy độ '
    'chính xác thấp.'
)
bullet_black(
    'Cẩn trọng khi convert sang RR (Risk Ratio): với prevalence outcome ~ 10–15%, '
    'RR thực có thể ~ 3–5 lần — không phải 11,5.'
)
bullet_black(
    'Khuyến nghị nội dung báo cáo: "Áp lực học tập rất cao có liên quan rất mạnh với '
    'lo âu nặng (OR ≥ 4 ở mức tin cậy 95%, ước lượng điểm OR=11,58, KTC 4,16–32,19) — '
    'cần kiểm chứng bằng cỡ mẫu lớn hơn."'
)

# =====================================================
# CÂU TRẢ LỜI (xanh, gom 1 chỗ trước phụ lục)
# =====================================================
H('6. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Tóm gọn:', bold=True)
para_blue(
    'CÓ — đủ mạnh để khẳng định "có liên quan rất mạnh", NHƯNG KHÔNG đủ chính xác '
    'để dùng OR=11,58 như con số tuyệt đối.'
)

para_blue('Phân tích chi tiết:', bold=True)

bullet_blue(
    'Tính ý nghĩa thống kê: KTC [4,16; 32,19] KHÔNG chứa 1 → CÓ ý nghĩa thống kê chắc '
    'chắn (p << 0,05). Đây là điểm chốt quan trọng nhất.'
)

bullet_blue(
    'Độ MẠNH (effect size): cận dưới = 4,16 → ngay cả ở tình huống "tệ nhất" trong '
    'khoảng tin cậy 95%, OR thực vẫn ≥ 4 lần. Theo phân loại Cohen/Hosmer-Lemeshow: '
    'OR > 4 = MẠNH; OR > 6 = RẤT MẠNH. Vậy effect size là MẠNH, không thể bác bỏ.'
)

bullet_blue(
    'Độ CHÍNH XÁC: width KTC = 32,19 − 4,16 = 28,03; tỷ số UL/LL = 7,74 → KTC RỘNG. '
    'OR thực có thể ở BẤT KỲ ĐÂU trong khoảng [4,16; 32,19] — không phải đúng 11,58. '
    'Đây là dấu hiệu cảnh báo về cỡ mẫu hoặc số ca outcome trong nhóm tham chiếu.'
)

bullet_blue(
    'Khả năng nguyên nhân KTC rộng: (a) cỡ mẫu Wen 2020 chỉ n=900 chia 6 mức Likert; '
    '(b) lo âu nặng trong nhóm "áp lực rất thấp" có thể chỉ vài % → ô tham chiếu nhỏ; '
    '(c) thầy nên check Bảng 3 bài gốc để xem n từng ô — nếu có ô < 10 quan sát thì '
    'OR phình to là chuyện đương nhiên.'
)

para_blue('Khuyến nghị cách trích dẫn vào báo cáo CTH:', bold=True)
bullet_blue(
    'KHÔNG nên: "Áp lực rất cao tăng nguy cơ lo âu nặng gấp 11,5 lần (Wen 2020)." → '
    'Sai cả về thuật ngữ ("nguy cơ" thay vì "odds") VÀ ngụy tạo chính xác.'
)
bullet_blue(
    'NÊN: "Áp lực học tập rất cao có liên quan rất mạnh với lo âu nặng so với áp lực '
    'rất thấp (OR=11,58; KTC 95% 4,16–32,19; Wen et al., 2020). Mức tăng odds tối '
    'thiểu là 4 lần (cận dưới); cần kiểm chứng bằng nghiên cứu cỡ mẫu lớn hơn vì KTC '
    'rộng cho thấy ước lượng điểm có độ chính xác thấp."'
)

para_blue('Bài học chung:', bold=True)
bullet_blue(
    'OR cao + KTC rộng rất phổ biến ở nghiên cứu cỡ mẫu vừa (n < 1000) với outcome '
    'phân nhóm. Người đọc cần luôn nhìn KTC trước khi tin vào con số ước lượng điểm.'
)
bullet_blue(
    'Quy tắc tay nhanh: nếu UL/LL > 4 → KTC rộng → effect size báo cáo cần coi như '
    '"hướng có thể xảy ra" hơn là "ước lượng chính xác".'
)
bullet_blue(
    'Trong meta-analysis của thầy sau này (nếu có), KTC như vậy sẽ được cho weight '
    'thấp do precision thấp → ảnh hưởng nhỏ đến pooled OR.'
)

# =====================================================
# PHỤ LỤC — REFERENCES
# =====================================================
H('7. Phụ lục — Tài liệu tham khảo', level=2)

para_black(
    '1. Wen, X., et al. (2020). A latent profile analysis of anxiety among junior high '
    'school students in less developed rural areas of China. International Journal of '
    'Environmental Research and Public Health. (QT008 trong DB dự án.)', italic=True
)

para_black(
    '2. Altman, D. G., & Bland, J. M. (1995). Absence of evidence is not evidence of '
    'absence. BMJ, 311(7003), 485. doi:10.1136/bmj.311.7003.485 — Cảnh báo phân biệt '
    '"không có ý nghĩa" vs "không có hiệu ứng" và vai trò của KTC.', italic=True
)

para_black(
    '3. Greenland, S., Senn, S. J., Rothman, K. J., Carlin, J. B., Poole, C., Goodman, '
    'S. N., & Altman, D. G. (2016). Statistical tests, P values, confidence intervals, '
    'and power: a guide to misinterpretations. European Journal of Epidemiology, 31(4), '
    '337–350. doi:10.1007/s10654-016-0149-3 — Hướng dẫn diễn giải KTC đúng cách.',
    italic=True
)

para_black(
    '4. Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). Applied Logistic '
    'Regression (3rd ed.). Hoboken, NJ: Wiley. — Sách kinh điển về OR và KTC trong '
    'logistic regression.', italic=True
)

para_black(
    '5. Zhang, J., & Yu, K. F. (1998). What\'s the relative risk? A method of correcting '
    'the odds ratio in cohort studies of common outcomes. JAMA, 280(19), 1690–1691. '
    '— Công thức quy đổi OR → RR khi outcome phổ biến (rất quan trọng cho OR=11,58).',
    italic=True
)

para_black('6. Tham khảo trong corpus dự án (DB Lo-au):')
bullet_black(
    'QT008 Wen et al. (2020) — bài thầy đang trích.', italic=True
)
bullet_black(
    'QT008 Qiu 2022 — báo cáo OR=6,74 (KTC 4,73–9,61) cho resilience↔depression — KTC '
    'hẹp hơn nhiều, đáng tin cậy hơn về precision.', italic=True
)
bullet_black(
    'QT014 Anderson 2025 — narrative review nhắc đến áp lực học tập là yếu tố chính '
    'tăng lo âu, nhưng không có pooled OR.', italic=True
)
bullet_black(
    'Doc đã viết: 01_Bao-cao/OR_Wen_2020_dien_dat_nguy_co_vs_odds.docx — về cách '
    'diễn đạt "nguy cơ" vs "odds" cho cùng OR này.', italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
