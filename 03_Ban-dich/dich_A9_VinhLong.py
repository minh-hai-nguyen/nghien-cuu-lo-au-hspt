# -*- coding: utf-8 -*-
"""Dịch A9 — Vĩnh Long 2024 — Bài VN → nhận xét + phản biện"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.51298/vmj.v541i1.10696', size=10)
add_heading(doc, 'Trầm cảm và một số yếu tố liên quan của học sinh THPT tại TP Vĩnh Long, tỉnh Vĩnh Long năm 2023', 1)
h = doc.add_paragraph()
r = h.add_run('Depression and Associated Factors among High School Students in Vinh Long City, Vinh Long Province in 2023')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề', 'Trầm cảm và một số yếu tố liên quan của HS THPT tại TP Vĩnh Long 2023'),
    ('Tác giả', 'Nguyễn Thanh Truyền (1), Đặng Phúc Vinh (2), Phạm Hồng Nhân (3), Nguyễn Minh Phương (4), Nguyễn Tấn Đạt (4)'),
    ('Cơ quan', '(1) BV đa khoa Vĩnh Long; (2) TT Kiểm soát bệnh tật tỉnh Vĩnh Long; (3) QNQD Healthcare Solutions; (4) ĐH Y Dược Cần Thơ'),
    ('Tạp chí', 'Tạp chí Y học Việt Nam, Tập 541, Số 1, tháng 8/2024, tr. 369–374'),
    ('DOI', '10.51298/vmj.v541i1.10696'),
    ('Loại NC', 'Cắt ngang mô tả có phân tích'),
    ('Mẫu', '919 HS THPT (lớp 10–12), 4 trường, TP Vĩnh Long, 10–12/2023'),
    ('Công cụ', 'CES-D (20 câu, 0–60; >21 triệu chứng, >25 trầm cảm) + ESSA (áp lực HT)'),
])

add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'TC Y học VN — có DOI, peer-reviewed VN.',
    'n = 919 — mẫu khá lớn cho NC VN cấp tỉnh.',
    'CES-D + ESSA — kết hợp đo trầm cảm + áp lực HT. ESSA cũng dùng trong VN21 (α=0,88) và UNICEF VN22.',
    'Trầm cảm 12,2% (CES-D >25) — THẤP hơn DASS-21 (40–86%) do CES-D ngưỡng cao hơn.',
    'Yếu tố gia đình + trường: sống cùng người nghiện rượu, bạo lực, thầy cô la mắng. Phù hợp UNICEF VN22.',
    'Vĩnh Long — Miền Tây, ít NC SKTT so với HN/TPHCM.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'CES-D đo TRẦM CẢM — không đo LO ÂU riêng. Đề tài cần GAD-7/DASS-21.',
    'Chỉ Vĩnh Long — không đại diện toàn quốc.',
    'TC Y học VN — không PubMed/Scopus.',
    'Bài không nêu OR cụ thể — chỉ p. Khó so sánh.',
]:
    add_p(doc, f'• {b}')

add_heading(doc, 'KẾT QUẢ CHÍNH', 2)
add_heading(doc, 'Bảng 1. Tỷ lệ trầm cảm và căng thẳng (n = 919)', 3)
add_table(doc,
    ['Chỉ số', 'Tỷ lệ', 'Ghi chú'],
    [['Triệu chứng trầm cảm (CES-D >21)', '8,8%', ''],
     ['Trầm cảm (CES-D >25)', '12,2%', 'Ngưỡng chẩn đoán'],
     ['Căng thẳng thấp (ESSA)', '42,8%', ''],
     ['Căng thẳng trung bình', '35,6%', ''],
     ['Căng thẳng cao', '21,7%', '1/5 HS áp lực cao']],
    widths=[5.0, 2.0, 4.0])

add_heading(doc, 'Bảng 2. Yếu tố nguy cơ trầm cảm (hồi quy đa biến)', 3)
add_table(doc,
    ['Yếu tố', 'Hướng', 'p'],
    [['Sống cùng người nghiện rượu', 'Tăng nguy cơ', '<0,05'],
     ['Sống cùng người trầm cảm/tâm thần', 'Tăng nguy cơ', '<0,05'],
     ['Bị văng tục chế giễu', 'Tăng nguy cơ', '<0,05'],
     ['Bị đánh đập', 'Tăng nguy cơ', '<0,05'],
     ['Tranh cãi gay gắt', 'Tăng nguy cơ', '<0,05'],
     ['Bị thầy cô la mắng hăm dọa', 'Tăng nguy cơ', '<0,05']],
    widths=[5.0, 3.0, 2.0])
add_p(doc, 'Lưu ý: Bài gốc không nêu OR cụ thể — chỉ nêu p < 0,05. Cần xem PDF gốc Bảng 4.', size=9, italic=True)

add_p(doc, 'Đối chiếu liên bài', bold=True)
add_p(doc, 'Trầm cảm 12,2% (CES-D) — thấp hơn DASS-21 (VN16 Bảo Quyên: 78,8%; VN18 An Giang: 47,3%) nhưng phù hợp V-NAMHS 2022 (VN02: chẩn đoán DISC-5 thấp). Yếu tố gia đình (bạo lực, nghiện rượu) phù hợp UNICEF VN22, Ngô Anh Vinh VN15 (ACEs→lo âu), Islam QT31 (bắt nạt AOR=1,68). Yếu tố thầy cô la mắng — phù hợp UNICEF VN22 (HS không thoải mái tìm GV) và Ireland QT32 (OGA bảo vệ). ESSA cũng dùng trong VN21 (β=4,73 học thêm).')

add_heading(doc, 'KẾT LUẬN', 2)
add_p(doc, 'Trầm cảm 12,2% (CES-D >25) ở 919 HS THPT Vĩnh Long. Yếu tố gia đình (nghiện rượu, bạo lực, chế giễu) và trường học (thầy cô la mắng) là nguy cơ chính. Mặc dù tỷ lệ trầm cảm trung bình cả nước, các yếu tố gia đình và trường học ảnh hưởng lớn — cần giải pháp phối hợp nhà trường + gia đình.')

add_abbreviation_table(doc, [
    ('CES-D', 'Center for Epidemiologic Studies Depression Scale — Thang Trầm cảm CES'),
    ('ESSA', 'Educational Stress Scale for Adolescents'),
    ('THPT', 'Trung học Phổ thông'),
    ('SKTT', 'Sức khỏe tâm thần'),
])

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'TC Y học VN có DOI (10.51298/vmj). n=919 khá lớn. Vĩnh Long — Miền Tây ít NC.',
    'CES-D + ESSA — kết hợp 2 thang đo. Hồi quy đa biến.',
    'Yếu tố gia đình (nghiện rượu, bạo lực) — phù hợp UNICEF VN22 và Ngô Anh Vinh VN15 (ACEs).',
    'Yếu tố thầy cô la mắng — phù hợp UNICEF VN22 (HS không thoải mái tìm GV), Ireland QT32 (OGA bảo vệ).',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'CES-D đo TRẦM CẢM — KHÔNG đo LO ÂU riêng. Đề tài cần GAD-7 hoặc DASS-21 cho lo âu.',
    'Chỉ TP Vĩnh Long — không đại diện toàn tỉnh/quốc.',
    'TC Y học VN — không lập chỉ mục PubMed/Scopus. Peer-review VN.',
    'Bài KHÔNG nêu OR cụ thể — chỉ p < 0,05. Rất khó so sánh.',
    'Không đo: screen time/MXH, giấc ngủ, bắt nạt mạng — yếu tố quan trọng.',
    'CES-D ngưỡng >25 (60 điểm) — khác DASS-21 ngưỡng. Không so sánh trực tiếp được.',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'Cần NC lo âu riêng (GAD-7/DASS-21) tại Miền Tây — chưa có.',
    'So sánh CES-D vs DASS-21 vs GAD-7 trên cùng mẫu HS Miền Tây.',
    'Thêm biến MXH, giấc ngủ, bắt nạt mạng vào mô hình.',
    'Mở rộng: các tỉnh Miền Tây khác (Cần Thơ, Đồng Tháp — UNICEF VN22 có Đồng Tháp).',
]:
    add_red(doc, f'• {s}')

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '40_VinhLong_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'A9 Vinh Long: {len(t)} chars, {len(d.tables)} tables')
