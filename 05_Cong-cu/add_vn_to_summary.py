# -*- coding: utf-8 -*-
"""Add Vietnamese papers to the summary doc"""
import os
from docx import Document
from docx.shared import Pt, RGBColor

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
doc_path = os.path.join(BASE, "DocFiles", "TOM_TAT_10_BAI_BO_SUNG.docx")
doc = Document(doc_path)

doc.add_page_break()
h = doc.add_heading("CAC BAI BAO VIET NAM VA DONG NAM A", level=1)
for run in h.runs:
    run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)

p = doc.add_paragraph()
run = p.add_run("5 bai bao tu Viet Nam lien quan truc tiep den chu de lo au va tram cam o hoc sinh. Danh so tu 12 tro di. PDF da tai: bai 12 va 14.")
run.font.size = Pt(11)

papers = [
    ("12", "V-NAMHS - Khao sat Quoc gia ve Suc khoe Tam than Thanh thieu nien Viet Nam",
     "UNICEF VN + DH Queensland + Bo LDTBXH", "2023", "Bao cao chinh phu (UNICEF/QCMHR)",
     "Nguon chinh thuc - Khao sat dai dien quoc gia dau tien cua VN",
     "N=9.781, dai dien quoc gia, DASS-21 phien ban tieng Viet da xac thuc.",
     "Tram cam nhe-nang: 10%. Lo au nhe-nang: 15.6%. 1/5 thanh thieu nien VN gap thach thuc SKTT.",
     "NGHIEN CUU DAI DIEN QUOC GIA DAU TIEN cua VN ve SKTT thanh thieu nien. Co mau lon nhat DNA (N=9.781). Su dung DASS-21 - so sanh truc tiep voi Saikia (24.4%), Nakie (66.7%), Bhardwaj (73.3%). Ty le VN thap hon nhieu (15.6%) - co the do phuong phap dai dien quoc gia vs mau thuan tien.",
     "Bao cao chinh phu, khong phai bai bao peer-reviewed. Phuong phap chi tiet can xem bao cao day du (153 trang).",
     "Nguon tham khao SO 1 cho bat ky nghien cuu nao ve SKTT thanh thieu nien tai VN."),

    ("13", "Prevalence and Factors Associated with Depression among High School Students in Ho Chi Minh City",
     "Tac gia Viet Nam (Truong DHYK TP.HCM)", "2022", "National J Community Medicine",
     "Tap chi quoc te, peer-reviewed",
     "N=384 (165 nam, 219 nu), THPT TP.HCM, CES-D.",
     "Tram cam: 57.6% (CES-D mean=17.3). Nu > nam. Ap luc hoc tap la yeu to chinh.",
     "Ty le 57.6% o TP.HCM >> V-NAMHS (10%). Khac biet do cong cu (CES-D vs DASS-21) va mau thuan tien vs dai dien QG. Cung CES-D voi Qiu (2022) - Qiu bao cao 26.0% o TQ.",
     "Co mau nho (N=384). Mot truong duy nhat.",
     "Du lieu tu do thi lon nhat VN. So sanh voi V-NAMHS va Qiu."),

    ("14", "Levels of Stress, Anxiety, Depression in Adolescents during/after COVID-19 in Vietnam",
     "Tac gia Viet Nam", "2023", "American J Psychiatric Rehabilitation",
     "Tap chi quoc te",
     "N=8.473, 6 tinh/thanh VN, DASS-21. So sanh TRONG vs SAU COVID-19.",
     "TRONG COVID: stress 65.5%, lo au 41.5%, tram cam 34.2%. SAU COVID: stress 55.4%, lo au 25.4%, tram cam 20.1%. Lo au giam 16 diem % sau dai dich.",
     "Nghien cuu DUY NHAT so sanh TRONG vs SAU COVID-19 o VN. Co mau lon (8.473). Lo au 41.5% trong COVID >> V-NAMHS 15.6%. So voi Xu (2021) 9.89% o TQ: VN cao hon nhieu.",
     "Tap chi it duoc biet den. Thiet ke cat ngang (hai mau khac nhau).",
     "So sanh COVID-19 voi Xu (2021). Cung DASS-21 voi V-NAMHS, Saikia, Nakie."),

    ("15", "Academic stress and depression among Vietnamese adolescents: moderated mediation model",
     "Tac gia VN (hop tac quoc te)", "2022", "Current Psychology, Springer",
     "Q2, IF~2.5, peer-reviewed, Scopus/WoS",
     "Mo hinh trung gian co dieu tiet: ap luc hoc tap - tram cam, vai tro cua hai long cuoc song va kha nang phuc hoi.",
     "Ap luc hoc tap lien quan thuan manh voi tram cam. Hai long cuoc song la trung gian. Phuc hoi la dieu tiet - hoc sinh phuc hoi cao bi anh huong it hon.",
     "Cung chu de resilience voi Qiu (2022). Qiu: phuc hoi thap OR=6.74 - nghien cuu VN xac nhan vai tro bao ve. Boi canh giao duc VN - ap luc thi cu rat cao. Tap chi Springer Q2.",
     "Khong bao cao ty le tram cam/lo au cu the. Thiet ke cat ngang.",
     "Xac nhan vai tro resilience o boi canh VN. So sanh voi Qiu (2022)."),

    ("16", "Mental health among ethnic minority adolescents in Vietnam and correlated factors",
     "Tac gia VN (hop tac quoc te)", "2024", "Mental Health & Prevention (Elsevier)",
     "Peer-reviewed, Elsevier",
     "Thanh thieu nien dan toc thieu so, truong noi tru VN.",
     "Stress: 24.7%. Lo au: 54.4%. Tram cam: 59.0%. Ty le rat cao so voi dan so chung VN.",
     "Lo au 54.4% >> V-NAMHS 15.6% = gap 3.5 lan. Tram cam 59.0% >> V-NAMHS 10% = gap 6 lan. BAT BINH DANG SKTT nghiem trong giua Kinh va thieu so. Cung boi canh thieu so/vung sau voi Wen (2020) TQ.",
     "Nhom dac thu (truong noi tru). Co mau khong ro.",
     "Bat binh dang SKTT tai VN. So sanh voi Saikia (thieu so Assam) va Wen (nong thon TQ)."),
]

for num, title, authors, year, journal, rank, method, results, highlights, limits, relevance in papers:
    doc.add_heading(f"Bai {num}: {title}", level=2)

    for label, value in [("Tac gia:", authors), ("Nam:", year), ("Tap chi:", journal),
                         ("Xep hang:", rank)]:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)

    for label, value, color in [
        ("Phuong phap:", method, None),
        ("Ket qua:", results, None),
        ("Diem noi bat:", highlights, RGBColor(0x00, 0x66, 0x00)),
        ("Han che:", limits, RGBColor(0xCC, 0x00, 0x00)),
        ("Lien quan 11 bai goc:", relevance, None),
    ]:
        p = doc.add_paragraph()
        run = p.add_run(label + " ")
        run.bold = True
        run.font.size = Pt(11)
        if color:
            run.font.color.rgb = color
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_paragraph()

doc.save(doc_path)
print("DONE: Added 5 Vietnamese papers")
