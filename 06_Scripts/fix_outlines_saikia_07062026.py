#!/usr/bin/env python
# -*- coding: utf-8 -*-
"""
fix_outlines_saikia_07062026.py
================================
Sửa fabrication N=287 (Saikia 2023) trong outlines bai-bao-Q1.
Paper gốc: Saikia AM, Das H, Rajendran V (2023). Indian J Community Med 48(6):835-840.
N THỰC TẾ = 360 (180 nam + 180 nữ, 10 trường x 36 học sinh, tuổi 14-17).
N=287 trong outline là BỊA — không có sub-sample 287 nào trong PDF gốc.

REVIEW-BEFORE-RUN: script này tạo bản FIXED, KHÔNG ghi đè file gốc.

Usage:
    python 06_Scripts/fix_outlines_saikia_07062026.py [--dry-run] [--apply]

Mặc định = dry-run (chỉ in patch). Thêm --apply để ghi file _FIXED_07062026.docx.
"""
from __future__ import annotations

import argparse
import io
import os
import re
import sys
from pathlib import Path

# Đảm bảo stdout/stderr in được tiếng Việt trên Windows console (cp1252).
if hasattr(sys.stdout, "buffer"):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8", errors="replace")
if hasattr(sys.stderr, "buffer"):
    sys.stderr = io.TextIOWrapper(sys.stderr.buffer, encoding="utf-8", errors="replace")

try:
    from docx import Document
except ImportError:
    sys.stderr.write("[ERROR] Cần cài python-docx: pip install python-docx\n")
    sys.exit(1)


# Repo root = thư mục cha của 06_Scripts/
ROOT = Path(__file__).resolve().parent.parent

# Danh sách outline KNOWN chứa N=287 (theo audit 07/06/2026).
# Các outline khác trong bai-bao-Q1/ có nhắc Saikia nhưng KHÔNG ghi N nên không cần sửa.
TARGETS = [
    ROOT / "bai-bao-Q1" / "OutlineBilingual_Q1_01062026.docx",
]

# Quy tắc thay thế (áp dụng trên text của từng paragraph/run).
# Mỗi tuple: (regex_pattern, replacement, mô tả).
REPLACEMENTS = [
    # 1. "trên 287 thanh thiếu niên" → "trên 360 thanh thiếu niên (180 nam/180 nữ)"
    (
        re.compile(r"trên\s+287\s+thanh\s+thiếu\s+niên", re.IGNORECASE),
        "trên 360 thanh thiếu niên (180 nam/180 nữ)",
        "VN: thay '287 thanh thiếu niên' → '360 thanh thiếu niên (180 nam/180 nữ)'",
    ),
    # 2. "on 287 adolescents" → "on 360 adolescents (180 boys/180 girls)"
    (
        re.compile(r"on\s+287\s+adolescents", re.IGNORECASE),
        "on 360 adolescents (180 boys/180 girls)",
        "EN: thay 'on 287 adolescents' → 'on 360 adolescents (180 boys/180 girls)'",
    ),
    # 3. "N=287" → "N=360"   (fallback nếu có)
    (
        re.compile(r"\bN\s*=\s*287\b"),
        "N=360",
        "FALLBACK: N=287 → N=360",
    ),
    # 4. "287 học sinh" → "360 học sinh (180 nam/180 nữ)"
    (
        re.compile(r"\b287\s+học\s+sinh\b", re.IGNORECASE),
        "360 học sinh (180 nam/180 nữ)",
        "FALLBACK VN: '287 học sinh' → '360 học sinh (180 nam/180 nữ)'",
    ),
    # 5. "287 students" → "360 students (180 boys/180 girls)"
    (
        re.compile(r"\b287\s+students\b", re.IGNORECASE),
        "360 students (180 boys/180 girls)",
        "FALLBACK EN: '287 students' → '360 students (180 boys/180 girls)'",
    ),
]


def fix_paragraph_runs(paragraph) -> list[tuple[str, str, str]]:
    """
    Áp dụng REPLACEMENTS trên TEXT của paragraph.
    Vì regex có thể vắt qua nhiều run, ta gộp text các run, replace,
    rồi nếu thay đổi thực sự xảy ra, set vào run đầu tiên và clear run còn lại.
    Trả về list (old_text, new_text, rule_desc) các thay đổi đã apply.
    """
    changes: list[tuple[str, str, str]] = []
    if not paragraph.runs:
        return changes

    full = "".join(r.text for r in paragraph.runs)
    new_full = full
    applied_rules: list[str] = []
    for pattern, replacement, desc in REPLACEMENTS:
        candidate = pattern.sub(replacement, new_full)
        if candidate != new_full:
            applied_rules.append(desc)
            new_full = candidate

    if new_full != full:
        changes.append((full, new_full, "; ".join(applied_rules)))
        # Ghi lại vào run đầu tiên, clear các run còn lại — giữ format của run đầu.
        first = paragraph.runs[0]
        first.text = new_full
        for r in paragraph.runs[1:]:
            r.text = ""
    return changes


def fix_document(doc) -> list[tuple[str, str, str, str]]:
    """Trả về list (location, old, new, rule)."""
    all_changes: list[tuple[str, str, str, str]] = []
    for idx, p in enumerate(doc.paragraphs):
        for old, new, rule in fix_paragraph_runs(p):
            all_changes.append((f"paragraph[{idx}]", old, new, rule))
    for ti, t in enumerate(doc.tables):
        for ri, row in enumerate(t.rows):
            for ci, cell in enumerate(row.cells):
                for pi, p in enumerate(cell.paragraphs):
                    for old, new, rule in fix_paragraph_runs(p):
                        loc = f"table[{ti}].row[{ri}].cell[{ci}].para[{pi}]"
                        all_changes.append((loc, old, new, rule))
    return all_changes


def strip_metadata(doc) -> None:
    """Xoá metadata (author, last_modified_by, comments) khỏi core properties."""
    cp = doc.core_properties
    try:
        cp.author = ""
        cp.last_modified_by = ""
        cp.comments = ""
        cp.title = ""
        cp.subject = ""
        cp.keywords = ""
        cp.category = ""
    except Exception as e:
        sys.stderr.write(f"[WARN] strip_metadata: {e}\n")


def process_file(path: Path, apply: bool) -> None:
    if not path.exists():
        print(f"[MISS] {path} không tồn tại — bỏ qua.")
        return
    print(f"\n=== {path.relative_to(ROOT)} ===")
    doc = Document(str(path))
    changes = fix_document(doc)
    if not changes:
        print("  (không tìm thấy N=287 / Saikia 287 — bỏ qua)")
        return
    for loc, old, new, rule in changes:
        print(f"  [{loc}] rule: {rule}")
        # In OLD/NEW gọn — cắt mỗi đoạn ~250 ký tự xung quanh thay đổi.
        print(f"    OLD: {old.strip()[:400]}")
        print(f"    NEW: {new.strip()[:400]}")
    if apply:
        strip_metadata(doc)
        out = path.with_name(path.stem + "_FIXED_07062026.docx")
        doc.save(str(out))
        print(f"  [SAVED] {out.relative_to(ROOT)}  (đã strip metadata)")
    else:
        print("  (dry-run — chưa ghi file. Dùng --apply để xuất _FIXED_07062026.docx)")


def main() -> int:
    ap = argparse.ArgumentParser(description=__doc__)
    ap.add_argument("--apply", action="store_true",
                    help="Thực sự ghi file _FIXED_07062026.docx (mặc định: dry-run).")
    ap.add_argument("--dry-run", action="store_true",
                    help="In patch không ghi file (mặc định).")
    args = ap.parse_args()
    apply = bool(args.apply) and not args.dry_run
    print(f"Mode: {'APPLY (ghi file)' if apply else 'DRY-RUN (chỉ in patch)'}")
    print(f"Targets ({len(TARGETS)} file):")
    for t in TARGETS:
        print(f"  - {t.relative_to(ROOT)}")
    for t in TARGETS:
        process_file(t, apply=apply)
    print("\nDone.")
    return 0


if __name__ == "__main__":
    sys.exit(main())
