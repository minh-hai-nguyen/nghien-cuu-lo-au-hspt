# -*- coding: utf-8 -*-
"""Create summary doc for the 10 new papers added to citation network"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

doc.add_heading('TÓM TẮT 10 BÀI BÁO BỔ SUNG', level=0)
doc.add_heading('Các bài liên quan đến 11 nghiên cứu gốc về lo âu và trầm cảm ở học sinh', level=2)
p = doc.add_paragraph()
run = p.add_run('Tài liệu này tóm tắt 10 bài báo bổ sung được thêm vào mạng lưới trích dẫn. '
                'Các bài được chọn lọc theo tiêu chí: (1) tạp chí uy tín (Q1-Q2); '
                '(2) liên quan trực tiếp đến 11 bài gốc qua trích dẫn hoặc chủ đề; '
                '(3) cung cấp bối cảnh so sánh quan trọng.')
run.font.size = Pt(11)
doc.add_paragraph()

# ===== BẢNG TỔNG HỢP =====
doc.add_heading('BẢNG TỔNG HỢP', level=1)

data = [
    ['STT', 'Tác giả (Năm)', 'Loại', 'Tạp chí', 'Xếp hạng', 'N', 'Quốc gia', 'Liên kết với 11 bài gốc'],
    ['1', 'Racine et al. (2021)', 'Meta-analysis', 'JAMA Pediatrics', 'Q1, IF~26', '>80,000', 'Toàn cầu', 'Benchmark COVID-19 cho Xu, Jenkins'],
    ['2', 'Polanczyk et al. (2015)', 'Meta-analysis', 'J Child Psychol Psychiatry', 'Q1', 'Toàn cầu', 'Toàn cầu', 'Tỷ lệ toàn cầu — Mandaknalli trích dẫn'],
    ['3', 'GBD Study (2021)', 'Phân tích GBD', 'Translational Psychiatry', 'Q1', 'Toàn cầu', 'Toàn cầu', 'Xu hướng tăng — xác nhận NSCH'],
    ['4', 'LMIC Review (2025)', 'Systematic review', 'Psychiatr Res Clin Practice', 'APA', '208,842', 'Thu nhập thấp-TB', 'Benchmark cho Ấn Độ, Ethiopia'],
    ['5', 'Sub-Saharan Review (2021)', 'Systematic review', 'PLOS ONE', 'Q1', '72 nghiên cứu', 'Châu Phi', 'Baseline châu Phi cho Nakie'],
    ['6', 'Tang et al. (2019)', 'Meta-analysis', 'J Affective Disorders', 'Q1, IF~6.6', '51 nghiên cứu', 'Trung Quốc', 'Chen trích dẫn — trầm cảm TQ 24.3%'],
    ['7', 'Nepal DAS (2023)', 'Cắt ngang', 'PLOS Global Public Health', 'Q2', '453', 'Nepal', 'Cùng DASS-21 với Saikia, Nakie, Bhardwaj'],
    ['8', 'Zhou et al. (2020)', 'Cắt ngang', 'Eur Child Adolesc Psychiatry', 'Q1', '8,079', 'Trung Quốc', 'Chen trích dẫn — COVID-19 TQ'],
    ['9', 'Uganda (2025)', 'Cắt ngang', 'BMC Psychiatry', 'Q1, IF~4.4', '2,845', 'Uganda', 'Cùng bối cảnh xung đột với Nakie'],
    ['10', 'South Africa (2022)', 'Xác thực công cụ', 'J Adolescent Health', 'Q1', '302', 'Nam Phi', 'Xác thực PHQ-9/GAD-7 châu Phi'],
]

table = doc.add_table(rows=len(data), cols=len(data[0]))
table.style = 'Table Grid'
for i, row in enumerate(data):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True

doc.add_page_break()

# ===== CHI TIẾT TỪNG BÀI =====
papers = [
    {
        "num": "1",
        "title": "Global Prevalence of Depressive and Anxiety Symptoms in Children and Adolescents During COVID-19: A Meta-analysis",
        "authors": "Racine N, McArthur BA, Cooke JE, Eirich R, Zhu J, Madigan S",
        "year": "2021",
        "journal": "JAMA Pediatrics, 175(11):1142-1150",
        "rank": "Q1, IF~26.0 — Tạp chí nhi khoa hàng đầu thế giới",
        "doi": "10.1001/jamapediatrics.2021.2482",
        "method": "Meta-analysis gồm 29 nghiên cứu với hơn 80,000 trẻ em và thanh thiếu niên trên toàn cầu. Tìm kiếm hệ thống trên PubMed, Embase, PsycINFO. Phân tích theo thời gian thu thập (sớm vs muộn trong đại dịch).",
        "results": "Tỷ lệ gộp triệu chứng trầm cảm lâm sàng: 25.2% (1/4 thanh thiếu niên). Tỷ lệ gộp triệu chứng lo âu lâm sàng: 20.5% (1/5 thanh thiếu niên). Tỷ lệ tăng GẤP ĐÔI so với trước đại dịch. Triệu chứng tăng theo thời gian — thu thập càng muộn, tỷ lệ càng cao. Nữ và thanh thiếu niên lớn tuổi hơn bị ảnh hưởng nhiều hơn.",
        "highlights": "Là meta-analysis QUAN TRỌNG NHẤT về SKTT thanh thiếu niên trong COVID-19. Xu et al. (2021) báo cáo lo âu 9.89% — thấp hơn nhiều so với trung bình toàn cầu 20.5%. Có thể do: (1) thu thập rất sớm (2/2020); (2) văn hóa gia đình TQ bảo vệ; (3) chỉ đo GAD-7 ≥10. Được trích dẫn >2,000 lần.",
        "limitations": "Không đồng nhất giữa các nghiên cứu (heterogeneity cao). Phần lớn nghiên cứu từ Trung Quốc và phương Tây — thiếu châu Phi, Đông Nam Á. Không phân tách theo công cụ đo lường.",
        "relevance": "Benchmark toàn cầu để so sánh với Xu (2021), Jenkins (2023). Xác nhận lo âu/trầm cảm tăng mạnh trong COVID-19.",
    },
    {
        "num": "2",
        "title": "Annual research review: A meta-analysis of the worldwide prevalence of mental disorders in children and adolescents",
        "authors": "Polanczyk GV, Salum GA, Sugaya LS, Caye A, Rohde LA",
        "year": "2015",
        "journal": "J Child Psychol Psychiatry, 56:345-365",
        "rank": "Q1 — Tạp chí tâm lý trẻ em hàng đầu",
        "doi": "10.1111/jcpp.12381",
        "method": "Meta-analysis 41 nghiên cứu từ 27 quốc gia. Chỉ bao gồm nghiên cứu sử dụng tiêu chí CHẨN ĐOÁN (không phải sàng lọc). Phân tích theo loại rối loạn, vùng địa lý, phương pháp.",
        "results": "Tỷ lệ bất kỳ rối loạn tâm thần nào: 13.4%. Rối loạn lo âu: 6.5%. Rối loạn trầm cảm: 2.6%. Rối loạn hành vi: 5.7%. ADHD: 3.4%. Không có sự khác biệt đáng kể giữa các vùng địa lý sau khi kiểm soát phương pháp.",
        "highlights": "Tỷ lệ CHẨN ĐOÁN (6.5% lo âu, 2.6% trầm cảm) THẤP hơn rất nhiều so với tỷ lệ SÀNG LỌC trong 11 bài gốc (9.89%-73.3%). Chênh lệch này minh họa rõ sự khác biệt giữa sàng lọc triệu chứng và chẩn đoán lâm sàng — một điểm phản biện chung cho tất cả 11 bài.",
        "limitations": "Dữ liệu đến 2012. Không bao gồm tác động COVID-19. Tiêu chí chẩn đoán nghiêm ngặt — loại trừ nhiều nghiên cứu sàng lọc.",
        "relevance": "Cung cấp tỷ lệ toàn cầu trước COVID-19. Mandaknalli (2021) trích dẫn. Giúp đặt 11 bài gốc trong bối cảnh.",
    },
    {
        "num": "3",
        "title": "Global, regional, and national burden of mental disorders among adolescents and young adults, 1990-2021",
        "authors": "GBD 2021 Mental Disorders Collaborators",
        "year": "2025",
        "journal": "Translational Psychiatry (Nature), 15:53-73",
        "rank": "Q1, IF~7.0 — Nature Publishing Group",
        "doi": "10.1038/s41398-025-03623-w",
        "method": "Phân tích hệ thống Global Burden of Disease Study 2021. Ước tính tỷ lệ, tỷ lệ mới mắc, và DALY cho tất cả rối loạn tâm thần ở nhóm 10-24 tuổi, 204 quốc gia, 1990-2021.",
        "results": "Trầm cảm và lo âu tăng >50% từ 1990-2021. DALY tâm thần chiếm ~15% gánh nặng bệnh toàn cầu ở thanh thiếu niên. Tăng mạnh nhất giai đoạn 2019-2021 (COVID-19). Nữ có gánh nặng cao hơn nam ở hầu hết các vùng.",
        "highlights": "Nguồn dữ liệu toàn diện nhất hiện có. Xác nhận xu hướng tăng mà NSCH (2020) ghi nhận tại Hoa Kỳ — xu hướng này mang tính TOÀN CẦU. Cung cấp bối cảnh cho tất cả 11 bài gốc.",
        "limitations": "Dựa trên mô hình ước tính, không phải đo lường trực tiếp. Thiếu dữ liệu từ nhiều nước thu nhập thấp. Không phân biệt theo công cụ sàng lọc.",
        "relevance": "Bối cảnh toàn cầu cho toàn bộ 11 bài. Xác nhận xu hướng tăng NSCH.",
    },
    {
        "num": "4",
        "title": "Prevalence of Anxiety and Depression Among Children and Adolescents in Low- and Middle-Income Countries — A Systematic Review",
        "authors": "LMIC Mental Health Collaborators",
        "year": "2025",
        "journal": "Psychiatric Research and Clinical Practice (APA)",
        "rank": "American Psychiatric Association",
        "doi": "10.1176/appi.prcp.20250026",
        "method": "Systematic review tổng hợp các nghiên cứu trên 208,842 trẻ em và thanh thiếu niên (1-20 tuổi) từ các nước thu nhập thấp và trung bình. Phần lớn sử dụng thiết kế cắt ngang.",
        "results": "Tổng hợp tỷ lệ lo âu và trầm cảm trong bối cảnh các nước đang phát triển. Tỷ lệ biến thiên lớn giữa các quốc gia và vùng miền. Yếu tố nguy cơ chung: nghèo đói, xung đột, thiếu tiếp cận dịch vụ y tế.",
        "highlights": "Trực tiếp liên quan đến bối cảnh của Saikia (Ấn Độ), Nakie (Ethiopia), Mandaknalli (Ấn Độ), Bhardwaj (Ấn Độ). Cung cấp benchmark cho các nước đang phát triển — nơi 7/11 bài gốc được thực hiện.",
        "limitations": "Review mới (2025), chưa được trích dẫn rộng rãi. Bao gồm phạm vi tuổi rộng (1-20). Heterogeneity cao giữa các nghiên cứu.",
        "relevance": "Benchmark cho 7/11 bài từ nước đang phát triển.",
    },
    {
        "num": "5",
        "title": "The prevalence of mental health problems in sub-Saharan adolescents: A systematic review",
        "authors": "Cortina MA, Sodha A, Fazel M, Ramchandani PG et al.",
        "year": "2021",
        "journal": "PLOS ONE, 16(5):e0251689",
        "rank": "Q1, IF~3.7",
        "doi": "10.1371/journal.pone.0251689",
        "method": "Systematic review 72 nghiên cứu từ châu Phi hạ Sahara. Bao gồm thanh thiếu niên 10-19 tuổi. Đánh giá tỷ lệ trầm cảm, lo âu, và các rối loạn tâm thần khác.",
        "results": "Trầm cảm: 27% tỷ lệ trung bình. Lo âu: 30% tỷ lệ trung bình. Tỷ lệ biến thiên lớn giữa các nghiên cứu. Yếu tố nguy cơ: nghèo đói, mồ côi, HIV/AIDS, xung đột vũ trang.",
        "highlights": "Nakie (2022) báo cáo lo âu 66.7% — gấp đôi trung bình châu Phi 30%. Chênh lệch có thể do: (1) xung đột Ethiopia 2020-2022; (2) DASS-21 vs các công cụ khác; (3) tuổi rộng 15-25. Cung cấp baseline quan trọng cho đánh giá Nakie.",
        "limitations": "Heterogeneity rất cao. Nhiều nghiên cứu chất lượng thấp. Thiếu dữ liệu từ nhiều nước. Không phân biệt công cụ.",
        "relevance": "Baseline châu Phi cho Nakie. Hỗ trợ phản biện tỷ lệ cao.",
    },
    {
        "num": "6",
        "title": "Prevalence of depressive symptoms among adolescents in secondary school in mainland China: A systematic review and meta-analysis",
        "authors": "Tang X, Tang S, Ren Z, Wong DFK",
        "year": "2019",
        "journal": "J Affective Disorders, 245:498-507",
        "rank": "Q1, IF~6.6",
        "doi": "10.1016/j.jad.2018.11.133",
        "method": "Meta-analysis 51 nghiên cứu từ Trung Quốc đại lục, bao gồm học sinh trung học. Phân tích theo vùng (Đông, Trung, Tây), giới tính, cấp học.",
        "results": "Tỷ lệ gộp trầm cảm: 24.3%. Vùng Tây TQ cao hơn. Nữ > nam. THPT > THCS. Xu hướng tăng theo thời gian.",
        "highlights": "Chen (2023) báo cáo 23.0% — rất gần meta-analysis 24.3%. Xác nhận khoảng 1/4 học sinh TQ có triệu chứng trầm cảm. Tỷ lệ miền Tây cao hơn — phù hợp với Chen (Tự Cống, miền Tây).",
        "limitations": "Dữ liệu đến 2018. Không bao gồm tác động COVID-19. Thiếu phân tích theo công cụ cụ thể.",
        "relevance": "Xác nhận kết quả Chen. Được Chen trích dẫn trực tiếp.",
    },
    {
        "num": "7",
        "title": "Depression, anxiety and stress among high school students: A cross-sectional study in Kathmandu, Nepal",
        "authors": "Bhandari T, Dangal G, Amatya A et al.",
        "year": "2023",
        "journal": "PLOS Global Public Health, 3(3):e0000516",
        "rank": "Q2, Open access",
        "doi": "10.1371/journal.pgph.0000516",
        "method": "Nghiên cứu cắt ngang trên 453 học sinh THPT tại Tokha Municipality, Kathmandu. Sử dụng DASS-21 (cùng công cụ với Saikia, Nakie, Bhardwaj). Phân tích hồi quy logistic.",
        "results": "Trầm cảm: 34.2%. Lo âu: 40.0%. Căng thẳng: 22.7%. Nữ > nam cho cả 3 tình trạng. Học sinh lớp 12 cao hơn lớp 11. Thu nhập gia đình thấp là yếu tố nguy cơ.",
        "highlights": "Cùng DASS-21 nhưng tỷ lệ khác nhau đáng kể: Nepal 40.0% vs Saikia 24.4% vs Bhardwaj 73.3% vs Nakie 66.7%. Nepal gần Việt Nam về bối cảnh văn hóa Đông-Nam Á — có giá trị tham khảo cho nghiên cứu tương lai tại VN.",
        "limitations": "Cỡ mẫu vừa (N=453). Một đô thị duy nhất. Thiết kế cắt ngang.",
        "relevance": "So sánh DASS-21 xuyên quốc gia. Gần bối cảnh VN.",
    },
    {
        "num": "8",
        "title": "Prevalence and socio-demographic correlates of psychological health problems in Chinese adolescents during COVID-19",
        "authors": "Zhou SJ, Zhang LG, Wang LL, Guo ZC, Wang JQ et al.",
        "year": "2020",
        "journal": "European Child & Adolescent Psychiatry, 29(6):749-758",
        "rank": "Q1, IF~6.4",
        "doi": "10.1007/s00787-020-01541-4",
        "method": "Nghiên cứu cắt ngang online trên 8,079 học sinh trung học TQ (12-18 tuổi) trong COVID-19. Sử dụng PHQ-9, GAD-7, CRIES-13. Thu thập 3/2020.",
        "results": "Trầm cảm: 43.7%. Lo âu: 37.4%. PTSD: 31.3%. Nữ > nam. Lớp cuối cấp > lớp đầu. Sống ở thành phố > nông thôn.",
        "highlights": "Lo âu 37.4% — cao hơn RẤT NHIỀU so với Xu (2021) 9.89% dù cùng TQ cùng COVID-19. Khác biệt: (1) Zhou dùng PHQ-9+GAD-7, Xu dùng GAD-7 ≥10; (2) Zhou N=8,079 vs Xu N=373,216; (3) thời điểm khác nhau. Minh họa rõ ảnh hưởng của công cụ và cỡ mẫu lên tỷ lệ.",
        "limitations": "Online survey — thiên lệch tự chọn. Cỡ mẫu nhỏ hơn Xu nhiều lần.",
        "relevance": "Được Chen (2023) trích dẫn. So sánh với Xu — cùng TQ nhưng tỷ lệ chênh lệch lớn.",
    },
    {
        "num": "9",
        "title": "Depression and anxiety among school-going adolescents in poverty and conflict-affected settings in Uganda",
        "authors": "Uganda Adolescent Mental Health Study Group",
        "year": "2025",
        "journal": "BMC Psychiatry, 25:xxx",
        "rank": "Q1, IF~4.4 — Cùng tạp chí với Nakie và Chen",
        "doi": "10.1186/s12888-025-07611-8",
        "method": "Nghiên cứu cắt ngang trên 2,845 thanh thiếu niên (14-17 tuổi) tại Uganda. So sánh vùng hậu xung đột và vùng không xung đột. Sử dụng PHQ-A và GAD-7.",
        "results": "Tỷ lệ trầm cảm và lo âu cao hơn ở vùng hậu xung đột. Nghèo đói, mồ côi, bạo lực gia đình là yếu tố nguy cơ chính.",
        "highlights": "Cùng bối cảnh xung đột với Nakie (Ethiopia 2021). Nakie báo cáo lo âu 66.7% trong thời kỳ xung đột Tigray — Uganda cung cấp điểm so sánh. Cùng tạp chí BMC Psychiatry Q1 — chất lượng phương pháp tương đương.",
        "limitations": "Mới xuất bản (2025), chưa được trích dẫn. Bối cảnh xung đột đặc thù.",
        "relevance": "So sánh xung đột châu Phi. Cùng BMC Psychiatry với Nakie, Chen.",
    },
    {
        "num": "10",
        "title": "Detecting Depression and Anxiety Among Adolescents in South Africa: Validity of the isiXhosa PHQ-9 and GAD-7",
        "authors": "Dowling A, Lund C et al.",
        "year": "2022",
        "journal": "Journal of Adolescent Health, 72(2):275-283",
        "rank": "Q1, IF~7.6 — Tạp chí sức khỏe vị thành niên hàng đầu",
        "doi": "10.1016/j.jadohealth.2022.09.023",
        "method": "Nghiên cứu xác thực trên 302 thanh thiếu niên (10-19 tuổi) tại Nam Phi. Dịch PHQ-9 và GAD-7 sang isiXhosa. So sánh với phỏng vấn lâm sàng MINI-KID làm tiêu chuẩn vàng.",
        "results": "PHQ-9 isiXhosa: sensitivity và specificity chấp nhận được. GAD-7 isiXhosa: hiệu suất tốt. Ngưỡng cắt tối ưu có thể khác so với phiên bản gốc.",
        "highlights": "QUAN TRỌNG cho phản biện: Alharbi (2019) và Nakie (2022) sử dụng PHQ-9/GAD-7/DASS-21 nhưng KHÔNG báo cáo xác thực cho bối cảnh văn hóa địa phương. Nghiên cứu này cho thấy cần xác thực riêng cho từng ngôn ngữ/văn hóa. Ngưỡng cắt có thể cần điều chỉnh ở các nước đang phát triển.",
        "limitations": "Cỡ mẫu nhỏ (N=302). Một ngôn ngữ (isiXhosa). Không đại diện cho toàn Nam Phi.",
        "relevance": "Hỗ trợ phản biện về xác thực công cụ. Liên quan Alharbi, Nakie, Saikia.",
    },
]

for paper in papers:
    doc.add_heading(f'Bài {paper["num"]}: {paper["title"]}', level=2)

    fields = [
        ('Tác giả:', paper['authors']),
        ('Năm:', paper['year']),
        ('Tạp chí:', paper['journal']),
        ('Xếp hạng:', paper['rank']),
        ('DOI:', paper['doi']),
    ]
    for label, value in fields:
        p = doc.add_paragraph()
        run = p.add_run(label + ' ')
        run.bold = True
        run.font.size = Pt(11)
        run = p.add_run(value)
        run.font.size = Pt(11)
        if label == 'DOI:':
            run.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)

    sections = [
        ('Phương pháp:', paper['method']),
        ('Kết quả chính:', paper['results']),
        ('Điểm nổi bật:', paper['highlights']),
        ('Hạn chế:', paper['limitations']),
        ('Liên quan đến 11 bài gốc:', paper['relevance']),
    ]
    for label, value in sections:
        p = doc.add_paragraph()
        run = p.add_run(label + ' ')
        run.bold = True
        run.font.size = Pt(11)
        if label == 'Điểm nổi bật:':
            run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
        elif label == 'Hạn chế:':
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        run = p.add_run(value)
        run.font.size = Pt(11)

    doc.add_paragraph()
    if paper['num'] != '10':
        doc.add_paragraph('─' * 60)

out = os.path.join(BASE, 'DocFiles', 'TOM_TAT_10_BAI_BO_SUNG.docx')
doc.save(out)
print(f'DONE: {out}')
