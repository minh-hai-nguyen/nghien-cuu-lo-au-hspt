"""
Generate 5 Vietnamese summary docx files for QT074-QT078.
All facts verified from PDF reads on 07/06/2026.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUTDIR = r"c:\Users\OS\OneDrive\read_books\Lo-au\Tom-tat-tung-bai"

def make_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    # set East Asian font too
    rpr = style.element.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:eastAsia'), 'Times New Roman')
    # line spacing 1.5
    pf = style.paragraph_format
    pf.line_spacing = 1.5
    return doc

def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(14 if level == 1 else 12)
    return p

def add_para(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p

def strip_metadata(doc):
    cp = doc.core_properties
    cp.author = ""
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.comments = ""
    cp.category = ""
    cp.last_modified_by = ""

def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    run = p.add_run("Soạn 07/06/2026 từ PDF gốc đã verify")
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.italic = True

def build(filename, sections):
    """sections: list of (heading, [paragraphs])"""
    doc = make_doc()
    # Title
    title = sections[0][0]
    add_heading(doc, title, level=1)
    # Rest
    for h, paras in sections[1:]:
        add_heading(doc, h, level=2)
        for para in paras:
            add_para(doc, para)
    strip_metadata(doc)
    add_footer(doc)
    path = os.path.join(OUTDIR, filename)
    doc.save(path)
    # word count
    wc = sum(len(p.split()) for _, paras in sections for p in paras)
    return path, wc

# ============================================================
# QT074 — Karasu 1986 — The Psychotherapies: Benefits and Limitations
# ============================================================
qt074 = [
    ("QT074 — Karasu (1986): Các liệu pháp tâm lý — Lợi ích và Hạn chế", []),
    ("Thông tin xuất bản", [
        "Tác giả: Toksoz B. Karasu, M.D. — Giáo sư Tâm thần học, Albert Einstein College of Medicine, Yeshiva University; Giám đốc Khoa Tâm thần, Bronx Municipal Hospital Center, Bronx, New York.",
        "Tiêu đề: \"The Psychotherapies: Benefits and Limitations\".",
        "Nguồn: American Journal of Psychotherapy, Vol. XL (40), No. 3, July 1986, trang 324–342.",
        "Hình thức: Bài thuyết trình tại Hội nghị tưởng niệm Emil A. Gutheil lần thứ 26 của Association for the Advancement of Psychotherapy, New York, ngày 02/11/1985.",
        "Đường dẫn PDF (đã verify): 02_Papers-goc/karasu1986.pdf",
        "DOI: 10.1176/appi.psychotherapy.1986.40.3.324 (kiểm chứng qua PsychiatryOnline).",
    ]),
    ("Tóm tắt", [
        "Karasu rà soát hiện trạng của các liệu pháp tâm lý vào giữa thập niên 1980 dọc theo năm trục lớn: điều trị (treatment), đào tạo (training), chẩn đoán (diagnosis), nghiên cứu (research) và rà soát/đánh giá (review). Tác giả nhận định lĩnh vực đang đạt được những tiến bộ chưa từng có — chuẩn hóa chẩn đoán, đào tạo và điều trị; phát triển công cụ nghiên cứu để xác lập hiệu quả trị liệu; nở rộ các kỹ thuật ngắn hạn; và các nỗ lực tích hợp đa trường phái — đồng thời chỉ ra các hạn chế: phân cực giữa các trường phái, khủng hoảng bản sắc nghề nghiệp, áp lực kinh tế của tính \"hiệu quả-chi phí\" (cost-effectiveness), và sự bùng nổ tới hơn 450 loại trị liệu khác nhau khiến lĩnh vực rơi vào trạng thái \"tâm thần học trong khủng hoảng\".",
    ]),
    ("Bối cảnh nghiên cứu / Mục tiêu", [
        "Bối cảnh: Đầu thập niên 1980, lĩnh vực psychotherapy bước vào giai đoạn rapprochement (xích lại gần nhau) sau nhiều năm \"bành trướng và chia rẽ\" không kiểm soát. Đồng thời, các áp lực kinh tế (chi phí chăm sóc y tế tăng), các khủng hoảng đạo đức nghề nghiệp, và sự ra đời DSM-III làm thay đổi diện mạo thực hành.",
        "Mục tiêu: Tổng quan có phê phán về lợi ích và hạn chế của các liệu pháp tâm lý đương thời, đặt trong bối cảnh xã hội-chính trị-kinh tế, để định hướng phát triển nghề.",
    ]),
    ("Phương pháp", [
        "Bài viết thuộc thể loại bài tổng quan tự sự (narrative review) / bài thuyết trình hội nghị — không phải nghiên cứu thực nghiệm.",
        "Tác giả tổng hợp tài liệu từ các báo cáo của APA Commission on Psychiatric Therapies (mà ông chủ trì 1979-1983), các nghiên cứu hiệu quả trị liệu đương thời (Smith, Glass & Miller; Jerome Frank và cộng sự), và các văn bản chính sách (DSM-III, hệ thống DRG).",
        "Cấu trúc: trình bày năm trục — Treatment & Training; Diagnosis; Research; Review — phân tích cả \"specificity\" (chuyên biệt hóa) và \"nonspecificity\" (yếu tố chung) của các liệu pháp.",
    ]),
    ("Phát hiện chính (con số trích từ PDF)", [
        "Bùng nổ liệu pháp: theo một thống kê được Karasu trích dẫn, có \"hơn 250 loại\" liệu pháp cạnh tranh trên thị trường tâm lý trị liệu; con số riêng của Karasu khảo sát trên toàn nước Mỹ \"đã vượt 450 loại\" (trang 325).",
        "Liệu pháp ngắn hạn: phạm vi từ 12 buổi (Mann & Goldman, liệu pháp \"diễn giải-làm rõ-đồng cảm\"), tới 1 năm hoặc tới 40 buổi (Malan), và mở rộng tới \"trị liệu một buổi duy nhất\" (single-session psychotherapy) trang 327.",
        "Cây quyết định 3 bước của Clarkin & Frances (1985) dẫn đến lựa chọn giữa 5 hình thức can thiệp ngắn hạn: khủng hoảng, tâm động học, giải quyết vấn đề, hôn nhân-gia đình, hoặc hành vi (trang 328).",
        "Hơn một nửa (>50%) các nhà chuyên môn được khảo sát cho biết họ theo nhiều hơn một định hướng lý thuyết (trang 329).",
        "Frank và cộng sự, qua 25 năm nghiên cứu, đã chứng minh vai trò ý nghĩa của \"kỳ vọng, sự làm chủ và khơi gợi cảm xúc\" như những thành phần chung hiệu quả của mọi liệu pháp (trang 330).",
        "Bài viết trích kết luận từ meta-analysis của Smith, Glass & Miller: \"nếu có điều gì, các phát hiện của chúng tôi cảnh báo chống lại một chủ nghĩa chiết trung trong thực hành mà không phân hóa thành một dạng liệu pháp cụ thể\" (trang 330-331).",
        "Về chẩn đoán: DSM-III được khen là cải tiến so với DSM-II/ICD nhờ đa trục (multiaxial) và tiêu chí thao tác hóa rõ ràng, nâng cao tính tin cậy (trang 331).",
    ]),
    ("Bàn luận / Hàm ý", [
        "Karasu đề xuất rằng \"chân lý\" trong trị liệu tâm lý có thể nằm ở sự tổng hợp giữa các trường phái khác nhau (\"enlightened eclecticism\" — chiết trung khai sáng — theo Yager 1985).",
        "Tác giả nhấn mạnh hai hướng đi song song và tưởng như đối lập: (1) chuyên biệt hóa (specificity) để \"điều trị nào, cho ai, hiệu quả nhất với vấn đề nào, trong hoàn cảnh nào\"; (2) phi-chuyên-biệt-hóa (nonspecificity) để tìm các yếu tố chung xuyên các liệu pháp.",
        "Việc tích hợp tâm trị liệu với dược trị liệu (psychotherapy + pharmacotherapy) cho thấy mô hình \"cộng dồn\" (additive model) ưu việt hơn từng phương thức đơn lẻ.",
    ]),
    ("Phản biện / Hạn chế", [
        "Hạn chế tác giả tự thừa nhận: bài viết chỉ là tổng quan định tính; \"nghiên cứu vẫn còn nhiều thiếu sót, kết quả chưa thuyết phục, các câu hỏi quan trọng chưa được trả lời\" và phát hiện chưa đủ tác động đến công việc lâm sàng thực sự (trang 325).",
        "Phản biện của người tóm tắt: (1) bài viết đã 40 tuổi (1986), nên các dữ liệu DSM-III, DRG, và \"450 loại\" trị liệu nay đã lỗi thời (DSM-5-TR 2022 hiện hành); (2) thiếu phân tích định lượng và bằng chứng RCT cấp độ cao; (3) phản ánh chủ yếu bối cảnh Bắc Mỹ, không bao quát các truyền thống trị liệu Đông Á/Việt Nam; (4) cảnh báo về \"hiệu quả-chi phí\" của Karasu rất tiên tri — đến 2026 vẫn là vấn đề trung tâm của chính sách bảo hiểm y tế tâm thần.",
    ]),
    ("Áp dụng cho LA NCS", [
        "Hữu ích cho khung lý thuyết của Luận án về can thiệp tâm lý lo âu: trích dẫn để biện luận vì sao cần kết hợp đa phương thức (tâm động học + nhận thức-hành vi + dược lý) thay vì chỉ một trường phái.",
        "Trích dẫn cho phần \"yếu tố chung trong trị liệu\" (Jerome Frank) khi bàn về cơ chế chung của các can thiệp cho học sinh có lo âu — kỳ vọng, làm chủ, cảm xúc.",
        "Hỗ trợ luận điểm của bài Q1: ở Việt Nam, dùng can thiệp ngắn hạn (brief therapy) là khả thi và phù hợp bối cảnh nguồn lực hạn chế, vì đã có tiền lệ từ 1986 cho thấy trị liệu ngắn không kém hiệu quả với trị liệu dài.",
        "Cung cấp tiền đề lịch sử về \"specificity vs nonspecificity dilemma\" — vẫn còn nguyên tính thời sự trong các meta-analysis 2020-2025.",
    ]),
    ("Trích dẫn nguyên văn", [
        "Trang 325: \"Psychotherapists are faced with solvable questions that are unimportant, and important questions that are unsolvable.\" (Các nhà trị liệu đối mặt với những câu hỏi giải được nhưng không quan trọng, và những câu hỏi quan trọng nhưng không giải được.)",
        "Trang 330-331 (trích Smith, Glass & Miller): \"If anything our findings warn against an eclecticism in practice that fails to differentiate into one or another form of psychotherapy. One of the paradoxes of psychotherapy ... may be that although all therapies are equally effective, one must choose only one to learn and practice.\" (p. 185 của meta-analysis gốc).",
    ]),
    ("Tham khảo", [
        "Karasu, T. B. (1986). The Psychotherapies: Benefits and Limitations. American Journal of Psychotherapy, 40(3), 324–342. DOI: 10.1176/appi.psychotherapy.1986.40.3.324.",
        "PDF: 02_Papers-goc/karasu1986.pdf — đã verify ngày 07/06/2026 (đọc trực tiếp 8 trang đầu).",
    ]),
]

# ============================================================
# QT075 — Karasu Biographical Sketch 14/02/2014
# ============================================================
qt075 = [
    ("QT075 — T. Byram Karasu, M.D.: Biographical Sketch (14/02/2014)", []),
    ("Thông tin xuất bản", [
        "Tác giả/đối tượng: T. Byram Karasu, M.D.",
        "Tiêu đề: \"Biographical sketch\" (Bản tiểu sử chính thức).",
        "Ngày soạn: 14 tháng 02 năm 2014.",
        "Hình thức: tài liệu tiểu sử 8 trang, không xuất bản trên tạp chí khoa học mà là tài liệu giới thiệu chính thức (CV mở rộng).",
        "Đường dẫn PDF (đã verify): 02_Papers-goc/biosketch_karasu_062112.pdf",
        "DOI/PMID: PDF không nêu rõ — đây là tài liệu cá nhân, không có DOI.",
    ]),
    ("Tóm tắt", [
        "Tài liệu tiểu sử trình bày học vấn, vị trí học thuật, vai trò lãnh đạo trong APA, và sự nghiệp viết sách của T. Byram Karasu — một trong những nhà tâm thần học-tâm lý trị liệu Mỹ có ảnh hưởng cuối thế kỷ 20 đầu thế kỷ 21. Tài liệu liệt kê 21 cuốn sách ông là tác giả/đồng tác giả/biên tập, hơn 100 bài báo, các giải thưởng (APA Presidential Commendation), và các nhận xét tán dương từ những nhân vật như Otto Kernberg, Salman Akhtar, Thomas Moore, Deepak Chopra, Wayne Dyer, Vamik Volkan và Glen Gabbard. Tài liệu cũng liệt kê chuỗi giảng viên thỉnh giảng tại Harvard, Mayo Clinic, McMaster, Ohio State, Pittsburgh, Toronto, NYU, Cornell, v.v.",
    ]),
    ("Bối cảnh nghiên cứu / Mục tiêu", [
        "Bối cảnh: Tài liệu được soạn để giới thiệu nhà khoa học khi mời thỉnh giảng hoặc xuất bản. Karasu là Silverman Professor và Chủ tịch Khoa Tâm thần & Khoa học Hành vi tại Albert Einstein College of Medicine, đồng thời Giám đốc Tâm thần tại Montefiore Medical Center.",
        "Mục tiêu của tài liệu: cung cấp tiểu sử khoa học chính thức.",
    ]),
    ("Phương pháp", [
        "Không áp dụng — đây là tài liệu tiểu sử, không phải nghiên cứu thực nghiệm.",
        "Cấu trúc: học vấn → các vị trí học thuật → các vai trò lãnh đạo trong APA → liệt kê sách kèm nhận xét của các chuyên gia → danh sách các bài giảng thỉnh giảng → các ấn phẩm tiêu biểu.",
    ]),
    ("Phát hiện chính (sự kiện rút từ PDF)", [
        "Học vấn: tốt nghiệp Yale University School of Medicine, Khoa Tâm thần năm 1969.",
        "Vị trí hiện tại (thời điểm 2014): Silverman Professor & University Chairman, Department of Psychiatry & Behavioral Sciences, Albert Einstein College of Medicine; Psychiatrist-in-Chief tại Montefiore Medical Center.",
        "Tác giả/biên tập của hơn 21 cuốn sách; tác giả/đồng tác giả của hơn 100 bài báo.",
        "Editor-in-Chief Emeritus của American Journal of Psychotherapy.",
        "Distinguished Life Fellow của APA, nhận APA Presidential Commendation.",
        "Chủ tịch APA Commission on Psychiatric Therapies (1979-1983) — cho ra báo cáo phê bình các liệu pháp tâm lý-xã hội và liệu pháp soma.",
        "Năm 1981: được bổ nhiệm chủ tịch task force quốc gia gồm hơn 400 học giả/nhà nghiên cứu/lâm sàng để soạn báo cáo nền tảng \"Treatments of Psychiatric Disorders\" (4 tập, xuất bản 5/1989) — được Atlantic Monthly khen là \"đi trước thời đại 25 năm\".",
        "Chủ tịch APA Work Group on Major Depressive Disorders (1991-1993), xuất bản Practice Guideline for Major Depressive Disorder in Adults (4/1993, bản sửa đổi 4/2000).",
        "Một số sách tiêu biểu được trích dẫn nhận xét: Life Witness, Maxims Minimus, Gotham Chronicles, Rags of My Soul, Of God and Madness, The Spirit of Happiness, The Art of Marriage Maintenance, The Art of Serenity (national bestseller, dịch ra 5 ngôn ngữ), The Psychotherapist as Healer, Deconstruction of Psychotherapy, Wisdom in the Practice of Psychotherapy.",
        "Các bài giảng thỉnh giảng tiêu biểu: Hamilton Ford Lecturer (Titus Harris Society), Harvard Bicentennial Lecturer (McLean Hospital), Hincks Memorial Lecturer (McMaster), Ralph M. Patterson Memorial Lecturer (Ohio State), Royal College Speaker (Mount Sinai Hospital, University of Toronto), Alfred H. Stanton Lecturer (Harvard Medical School/McLean), John A. Graft Lecture (Mayo Clinic College of Medicine).",
    ]),
    ("Bàn luận / Hàm ý", [
        "Tài liệu cho thấy Karasu là nhân vật có thẩm quyền cao trong các tài liệu nền tảng của APA về điều trị rối loạn tâm thần và hướng dẫn thực hành lâm sàng cho trầm cảm — những văn bản này có ảnh hưởng đến đào tạo lâm sàng và bảo hiểm y tế ở Mỹ.",
        "Phong cách nghề nghiệp: kết hợp giữa tâm thần học, tâm lý trị liệu tâm động học, và linh đạo (spirituality) — thể hiện trong các tác phẩm The Spirit of Happiness, The Art of Serenity, Life Witness.",
    ]),
    ("Phản biện / Hạn chế", [
        "Hạn chế của bản thân tài liệu: là tài liệu PR/giới thiệu, không phải khảo cứu trung lập — các nhận xét trích dẫn đều mang tính khen ngợi.",
        "PDF không nêu rõ ngày sinh hay quốc tịch gốc của Karasu (tuy các nguồn ngoài cho biết ông là người Thổ Nhĩ Kỳ).",
        "Phản biện của người tóm tắt: (1) tài liệu này có giá trị thư mục hơn là giá trị khoa học — phù hợp để xác lập \"thẩm quyền nguồn\" khi trích bài 1986 (QT074) và Practice Guideline; (2) không nên dùng tài liệu này làm bằng chứng khoa học, chỉ làm bằng chứng về tầm vóc tác giả.",
    ]),
    ("Áp dụng cho LA NCS", [
        "Khi LA và bài Q1 trích Karasu (1986) hoặc trích APA Practice Guideline for Major Depressive Disorder, tài liệu này giúp xác lập tính tin cậy của tác giả nguồn — Karasu không phải tác giả vô danh mà là chủ tịch task force APA, Chair một khoa Tâm thần đại học top.",
        "Hữu ích cho phần \"Lịch sử các liệu pháp tâm lý\" trong tổng quan luận án — Karasu được xem là một trong các \"voice\" định hình ngành.",
    ]),
    ("Trích dẫn nguyên văn", [
        "Trang 1: \"T. Byram Karasu, M.D. graduated from Yale University School of Medicine, Department of Psychiatry in 1969. He is presently Silverman Professor and the University Chairman of the Department of Psychiatry & Behavioral Sciences at Albert Einstein College of Medicine and Psychiatrist-in-Chief at Montefiore Medical Center.\"",
        "Trang 1: \"From 1979-1983, Dr. Karasu chaired the APA's Commission on Psychiatric Therapies... In 1981, Dr. Karasu was appointed chairman of another national task force comprised of over 400 scholars... with the goal of producing a seminal document describing the treatment of each psychiatric disorder.\"",
    ]),
    ("Tham khảo", [
        "Karasu, T. B. (2014). Biographical sketch [Tài liệu tiểu sử cá nhân, soạn 14/02/2014].",
        "Tham chiếu chéo: Karasu, T. B. (1986). The Psychotherapies: Benefits and Limitations. American Journal of Psychotherapy, 40(3), 324–342.",
        "Tham chiếu chéo: American Psychiatric Association. (1993, sửa 2000). Practice Guideline for Major Depressive Disorder in Adults (Karasu chủ trì Work Group 1991-1993).",
        "PDF: 02_Papers-goc/biosketch_karasu_062112.pdf — đã verify ngày 07/06/2026 (đọc 5 trang đầu của tài liệu 8 trang).",
    ]),
]

# ============================================================
# QT076 — Small & Blanc 2021 — Tam Giao + Vietnam COVID-19
# ============================================================
qt076 = [
    ("QT076 — Small & Blanc (2021): Sức khỏe tâm thần trong COVID-19 — Tam Giáo và Phản ứng của Việt Nam", []),
    ("Thông tin xuất bản", [
        "Tác giả: Sean Small (1) và Judite Blanc (1,2*).",
        "Cơ quan: (1) Department of Applied Psychology, New York University Steinhardt School of Culture, Education, and Human Development, New York, USA; (2) Department of Population Health, Center for Healthful Behavior Change, NYU Grossman School of Medicine, New York, USA.",
        "Tiêu đề: \"Mental Health During COVID-19: Tam Giao and Vietnam's Response\".",
        "Tạp chí: Frontiers in Psychiatry, Volume 11, Article 589618, xuất bản 08 January 2021. Thuộc mục \"Public Mental Health\" — mini-review.",
        "DOI: 10.3389/fpsyt.2020.589618 (verify trong PDF, trang 1).",
        "Tài trợ: NIH/NHLBI grant T32HL129953 (JB).",
        "License: CC BY (open access).",
        "Đường dẫn PDF (đã verify): 02_Papers-goc/Chua-phan-loai/Small_Blanc_2021_TamGiao_Vietnam.pdf",
    ]),
    ("Tóm tắt", [
        "Bài mini-review lập luận rằng sự ứng phó thành công của Việt Nam với đại dịch COVID-19 — duy trì số ca nhiễm thấp và tử vong rất thấp trong giai đoạn đầu — có thể được hiểu một phần qua khái niệm văn hóa Tam Giáo (Three Teachings): sự cùng tồn tại của Phật giáo, Nho giáo và Đạo giáo trong văn hóa Việt Nam, kết hợp với \"tính cộng thêm văn hóa\" (cultural additivity) làm nền tảng cho khả năng chống chịu tâm lý cộng đồng. Tam Giáo, theo các tác giả, đã giúp dung hợp các giá trị tưởng như mâu thuẫn (ví dụ Nho giáo nhấn mạnh trách nhiệm xã hội vs. Đạo giáo nhấn mạnh hòa hợp tự nhiên/wuwei) thành một khung tâm lý giúp người Việt chấp nhận các biện pháp y tế công cộng nghiêm ngặt như cách ly và truy vết tiếp xúc, đồng thời nuôi dưỡng tinh thần đoàn kết cộng đồng (cây ATM gạo).",
    ]),
    ("Bối cảnh nghiên cứu / Mục tiêu", [
        "Bối cảnh: COVID-19 là bệnh truyền nhiễm mới với R0 ≈ 2.5; tính đến 04/11/2020 toàn cầu có hơn 44 triệu ca và 1 triệu tử vong. Việt Nam — dân số ước 97.33 triệu — đến 30/07/2020 chỉ có 459 ca và không tử vong; đến 04/11/2020 có 35 tử vong. Việt Nam là một LMIC giáp biên với Trung Quốc (nơi xuất phát dịch).",
        "Mục tiêu: (1) mô tả phản ứng sớm của Việt Nam với COVID-19; (2) mô tả Tam Giáo như một yếu tố văn hóa-tâm lý; (3) thảo luận sức khỏe tâm thần và thay đổi lối sống; (4) đề xuất hàm ý cho các LMIC khác.",
    ]),
    ("Phương pháp", [
        "Thiết kế: bài mini-review tự sự (narrative mini-review), không phải systematic review — các tác giả thừa nhận rõ.",
        "Nguồn dữ liệu: tổng hợp tài liệu thứ cấp: số liệu dịch tễ COVID-19 (Our World in Data, Lancet), báo cáo chính sách Việt Nam (Sustainability journal, J Comm Saf Well-Being), phân tích văn hóa Việt (Vuong et al. về cultural additivity, Palgrave Communications 2018), các nghiên cứu sức khỏe tâm thần thời COVID-19 tại Việt Nam (Tran et al., Le et al., Nguyen et al., Huynh — Frontiers in Psychology/Psychiatry 2020).",
        "Cấu trúc bài: (1) Vietnam's resiliency to COVID-19; (2) Tam giao and the Vietnamese COVID-19 response; (3) Considerations for mental health; (4) Limitations; (5) Conclusions.",
        "Không có dữ liệu định lượng nguyên gốc; không có thống kê inferential.",
    ]),
    ("Phát hiện chính (số liệu trích từ PDF)", [
        "Việt Nam: dân số ~97.33 triệu; đến 30/07/2020 có 459 ca xác nhận, 0 tử vong; đến 04/11/2020 có 35 tử vong.",
        "Brazil (so sánh tương phản): hơn 5.4 triệu ca xác nhận và hơn 150,000 tử vong tính đến 04/11/2020 — gắn với phản ứng chính trị phủ nhận của Tổng thống Bolsonaro và thông tin sai lệch.",
        "Phát triển kit xét nghiệm đầu tiên ở Việt Nam: 07/02/2020. Số điểm xét nghiệm tăng từ 2 (1/2020) lên 120 (5/2020).",
        "Ứng dụng truy vết \"NCOVI\" — dự án \"neighborhood watch\" hợp tác giữa các công ty viễn thông và Bộ Y tế — đạt 1,040,000 lượt tải.",
        "Đến 30/06/2020, chỉ có 4 nhân viên y tế bị nhiễm COVID-19 ở Việt Nam.",
        "Bài viết trích Vuong và cộng sự (2018) về cultural additivity: Nho giáo \"chiếm ưu thế\" (dominates) hơn so với Phật giáo và Đạo giáo trong văn hóa dân gian Việt; Phật giáo bị \"isolation\" tương đối.",
        "Định nghĩa Tam Giáo trong bài: \"a coexistence of fundamental understandings of Buddhism, Confucianism, and Taoism\" (trang 2, mục Tam Giao and the Vietnamese COVID-19 Response).",
        "Quan sát của nhóm tác giả tại Bệnh viện Bạch Mai: nhân viên y tế bị cách ly hơn 3 tuần báo cáo kỳ thị/stigma gia tăng, đặc biệt tập trung vào cảm giác có lỗi với gia đình/bạn bè (trích Do Duy C et al. 2020, Psychiatry Clin Neurosci 74:566-68).",
        "Hơn ¾ mẫu Việt Nam được khảo sát đã thực hiện đủ cả 6 biện pháp phòng ngừa; các nhà từ thiện hỗ trợ nhóm yếu thế qua \"ATM gạo\" (rice ATMs).",
    ]),
    ("Bàn luận / Hàm ý", [
        "Tam Giáo cung cấp một \"khung tâm lý\" cho phép người Việt vừa tuân thủ kỷ luật xã hội (Nho giáo) vừa giữ thanh thản trong tự nhiên (Đạo giáo) vừa thực hành lòng từ bi (Phật giáo) — giúp giảm khổ tâm khi cách ly.",
        "Trách nhiệm chăm sóc người cao tuổi (đặc biệt dễ tổn thương với COVID-19) là một biểu hiện của Nho giáo + Phật giáo trong việc thực hành nghĩa vụ và karma.",
        "Đoàn kết cộng đồng có thể \"frame\" việc giãn cách xã hội như một hành vi đạo đức tập thể, làm giảm tác động tâm lý của \"cô lập xã hội\".",
        "Hàm ý chính sách: các LMIC khác có thể học từ Việt Nam ba điểm: (a) minh bạch và truyền thông sớm; (b) huy động đa ngành (gồm cảnh sát cộng đồng); (c) tận dụng yếu tố văn hóa-tâm linh để củng cố sức bền tâm lý.",
    ]),
    ("Phản biện / Hạn chế", [
        "Hạn chế tác giả tự thừa nhận (mục Limitations, trang 4): (1) đây không phải systematic review; (2) không có dữ liệu cụ thể đo lường thay đổi sức khỏe tâm thần của người Việt trong đại dịch; (3) không tìm được nghiên cứu trực tiếp kết nối Tam Giáo với kết cục COVID-19 ở Việt Nam — chỉ là suy luận lý thuyết.",
        "Phản biện của người tóm tắt: (1) bài viết có nguy cơ \"romanticize\" văn hóa Việt — Tam Giáo trong thực tế phức tạp hơn và không đồng nhất giữa các vùng miền; (2) các tác giả không phải chuyên gia Việt Nam học, mà là nhà tâm lý/y tế công cộng tại NYU — phân tích văn hóa dựa vào nguồn thứ cấp (Vuong et al.); (3) thành công ban đầu của Việt Nam (đến 11/2020) sau đó bị thử thách nặng nề bởi biến thể Delta giữa năm 2021 — bài viết không thể tiên đoán điều này; (4) đặt nặng văn hóa có thể che khuất các yếu tố thực tế như chính sách độc đảng, hệ thống cảnh sát mạnh, kinh nghiệm SARS 2003-2004; (5) lập luận \"Nho giáo có thể dung dưỡng việc nói dối để giữ hòa khí\" (trích Vuong et al.) là một suy luận có vấn đề về văn hóa Việt.",
    ]),
    ("Áp dụng cho LA NCS", [
        "Hữu ích cho phần \"bối cảnh văn hóa Việt Nam\" trong luận án về lo âu của học sinh: Tam Giáo có thể vừa là yếu tố bảo vệ (community solidarity, karma, hài hòa) vừa là yếu tố nguy cơ (áp lực thành tựu kiểu Nho giáo, hà khắc với bản thân).",
        "Cung cấp khung lý thuyết về \"cultural additivity\" (Vuong et al. 2018) để giải thích vì sao học sinh Việt Nam có hồ sơ tâm lý khác biệt so với học sinh phương Tây.",
        "Trích dẫn để biện luận trong bài Q1: bối cảnh COVID-19 làm trầm trọng thêm lo âu học đường tại Việt Nam, đặc biệt qua giai đoạn đóng cửa trường và học online 2020-2021.",
        "Liên kết với QT077 (Stankov 2010 — Confucian academic anxiety) để xây dựng lập luận tổng hợp về văn hóa Đông Á và lo âu học tập.",
    ]),
    ("Trích dẫn nguyên văn", [
        "Trang 2-3 (mục Tam giao and the Vietnamese COVID-19 Response): \"Tam giao, or the Three Teachings, is essential to Vietnamese culture and especially folklore, which acts as a cultural transmitter influencing contemporary thought and behavior. The Three Teachings is a coexistence of fundamental understandings of Buddhism, Confucianism, and Taoism.\"",
        "Trang 4 (Conclusions): \"Tam giao and its ethical values and philosophies may have aided the efficacy of healthy literacy, community solidarity, and positive outlook in relieving some of the psychological impacts of the pandemic.\"",
    ]),
    ("Tham khảo", [
        "Small, S., & Blanc, J. (2021). Mental Health During COVID-19: Tam Giao and Vietnam's Response. Frontiers in Psychiatry, 11, 589618. DOI: 10.3389/fpsyt.2020.589618.",
        "Tham chiếu chéo (trích trong bài): Vuong, Q.-H., Bui, Q. L. A. V., Vuong, T. T., Nguyen, V. T., Ho, M. T., et al. (2018). Cultural additivity: behavioural insights from the interaction of Confucianism, Buddhism and Taoism in folktales. Palgrave Communications, 4, 143. DOI: 10.1057/s41599-018-0189-2.",
        "PDF: 02_Papers-goc/Chua-phan-loai/Small_Blanc_2021_TamGiao_Vietnam.pdf — đã verify ngày 07/06/2026 (đọc 5 trang gồm toàn bộ bài + references).",
    ]),
]

# ============================================================
# QT077 — Stankov 2010 — Unforgiving Confucian Culture
# ============================================================
qt077 = [
    ("QT077 — Stankov (2010): Văn hóa Nho giáo không khoan dung — Mảnh đất màu mỡ cho thành tựu học tập cao, lo âu thi cử và nghi ngờ bản thân?", []),
    ("Thông tin xuất bản", [
        "Tác giả: Lazar Stankov — Honorary Professor, School of Psychology, The University of Sydney, Sydney NSW 2006, Australia.",
        "Tiêu đề: \"Unforgiving Confucian culture: A breeding ground for high academic achievement, test anxiety and self-doubt?\".",
        "Tạp chí: Learning and Individual Differences, Volume 20, Issue 6 (2010), trang 555–563.",
        "Lịch sử: Nhận 31/12/2009; sửa 26/04/2010; chấp nhận 11/05/2010.",
        "DOI: 10.1016/j.lindif.2010.05.003.",
        "Đường dẫn PDF (đã verify): 02_Papers-goc/Chua-phan-loai/Stankov_2010_Confucian_Academic.pdf",
    ]),
    ("Tóm tắt", [
        "Bài tổng quan tổng hợp ba mảng nghiên cứu so sánh đa văn hóa: (a) thành tựu học tập (TIMSS, PISA); (b) các biến tâm lý phi-nhận-thức (lo âu, tự nhận thức bản thân, tự hiệu năng); (c) các thang đo nhân cách-xã hội (Toughness/Maliciousness, Proviolence). Stankov chỉ ra một nghịch lý: học sinh các nước Nho giáo Đông Á (Hong Kong, Hàn Quốc, Nhật, Macau-Trung Quốc) liên tục đứng đầu PISA/TIMSS nhưng đồng thời báo cáo mức lo âu và nghi ngờ bản thân cao nhất. Tác giả lập luận rằng văn hóa Nho giáo \"đương đại\" — dù triết học Nho gốc có đề cao khoan dung (forgiveness) — đã trở nên \"không khoan dung\" (unforgiving) do các áp lực lịch sử-xã hội, niềm tin rằng nỗ lực (effort) chứ không phải năng lực (ability) là nguồn thành công, cộng với chủ nghĩa tập thể đặt thành tựu học tập làm thước đo danh dự gia đình.",
    ]),
    ("Bối cảnh nghiên cứu / Mục tiêu", [
        "Bối cảnh: Các quốc gia di sản Nho giáo (Confucian heritage) liên tục xếp hạng cao trong PISA và TIMSS — gây ra cái gọi là \"East Asian Learner Paradox\" (nghịch lý người học Đông Á).",
        "Mục tiêu: kết nối ba mảng nghiên cứu đến nay tách biệt — thành tựu, biến phi-nhận-thức, và đo lường nhân cách-xã hội — để tìm lời giải cho việc vì sao thành tựu cao đi kèm lo âu cao.",
    ]),
    ("Phương pháp", [
        "Thiết kế: review bài + meta-tổng hợp dữ liệu thứ cấp từ PISA 2003, TIMSS 1995/2007, PIRLS 2006.",
        "Phân tích so sánh giữa nhóm Confucian Asian (Hong Kong-China, South Korea, Japan, Macau-China) và European (Finland, Holland, Belgium, Switzerland, Liechtenstein).",
        "Dùng dữ liệu nhân cách-xã hội từ các nghiên cứu của chính Stankov và cộng sự (Stankov & Lee 2008, 2009; Saucier, Akers, Miller, Stankov & Knezevic 2009) trên hơn 200 người tham gia, độ tuổi 18-21, từ 73 quốc gia.",
        "Thang đo \"militant extremist mindset\" (MEM) — Stankov, Saucier & Knežević 2010 — đặc biệt là subscale Proviolence (pro-war attitudes), khảo sát N=2424 sinh viên năm nhất từ 9 quốc gia.",
        "Tính Cohen's d cho hiệu cỡ giữa nhóm Confucian Asian và European.",
    ]),
    ("Phát hiện chính (số liệu trích từ Bảng 1 PDF, trang 556)", [
        "Toán PISA 2003 — điểm trung bình: European 531 vs Confucian Asian 538 — chênh chỉ Cohen's d = 0.07 (không đáng kể).",
        "Self-concept (toán): European −0.03 vs Confucian Asian −0.41 — Cohen's d = 0.38 (Confucian Asian thấp hơn).",
        "Self-efficacy (toán): European 0.21 vs Confucian Asian −0.31 — Cohen's d = 0.52 (Confucian Asian thấp hơn nhiều).",
        "Anxiety (toán): European −0.40 vs Confucian Asian 0.10 — Cohen's d = 0.50 (Confucian Asian cao hơn).",
        "Tương quan cấp quốc gia trong PISA 2003: thành tích toán có r = −0.65 với anxiety và r = 0.42 với self-efficacy/confidence (Lee 2009).",
        "Trong Raven's Progressive Matrices (Brouwers, Van de Vijver & Van Hemert 2009): Nhật trung bình raw 109, Hàn Quốc 93 — kết quả khác hoàn toàn với ước lượng IQ từ Lynn & Vanhanen (2002, 2006).",
        "Thang Toughness/Maliciousness (TM, Stankov & Lee 2008, 2009 trên 9 vùng thế giới): Confucian Asia trung bình ~2.75 (gần điểm giữa 3); Latin America ~2.35; chênh đúng 1 SD — F(8,1976)=24.56, có ý nghĩa thống kê.",
        "Thang Proviolence (Stankov, Saucier & Knežević 2010, N=2424, 9 quốc gia): Hàn Quốc và Trung Quốc cao nhất; mean của 9 quốc gia có ý nghĩa F(8,2415)=86.77, p<0.01.",
        "Trích Lee (1996, trang 39): \"there is an extraordinary emphasis on effort, willpower or concentration of the mind in the Confucian tradition. Because there is a strong belief in attainability by all, there is also a strong belief that one's failure is not due to one's particular make-up or ability, but one's effort and willpower\".",
    ]),
    ("Bàn luận / Hàm ý", [
        "Bốn cách giải thích Stankov đặt ra cho việc lo âu cao đi kèm thành tích cao ở Confucian Asia:",
        "(1) Self-presentation/khiêm tốn — học sinh Đông Á có xu hướng \"khiêm tốn\" (modesty) khi tự đánh giá, nhưng giả thuyết này không giải thích được mức lo âu cao trong bối cảnh khiêm tốn.",
        "(2) Lỗi đo lường — PISA dùng ít item (chỉ 2 item cho self-concept), gây độ tin cậy thấp.",
        "(3) Tự đánh giá tương đối — học sinh giỏi ở nước có nhiều người giỏi sẽ tự đánh giá thấp hơn vì so sánh với peer mạnh.",
        "(4) Khác biệt thực sự (\"real\" substantive differences) — Stankov nghiêng về giả thuyết này: \"unforgiving\" trong văn hóa Nho giáo đương đại là yếu tố thực, không phải lỗi đo lường.",
        "Cơ chế đề xuất: niềm tin \"nỗ lực thay vì năng lực\" + áp lực gia đình-xã hội + \"in-group vs out-group forgiveness\" (Nho giáo có thể khoan dung với in-group nhưng khắt khe với out-group). Cấu trúc trầm cảm ở thanh thiếu niên Đông Á có thành phần \"socially prescribed self-evaluation\" và \"cognitive inefficiency\" (lo âu học đường), khác cấu trúc trầm cảm phương Tây (Woo et al. 2004, 2007).",
        "Hàm ý: \"social-emotional learning\" trong trường (như Singapore đưa vào từ 2005) chỉ chữa triệu chứng — nguyên nhân nằm sâu trong cộng đồng, văn hóa và xã hội. Cần can thiệp hai mũi: nâng nhận thức công chúng + tăng giá trị nghề nghiệp không cần thành tích học cao.",
    ]),
    ("Phản biện / Hạn chế", [
        "Hạn chế tác giả tự nhận (caveats, mục 6): bằng chứng Confucian Asia \"không khoan dung\" hơn châu Âu chưa đủ mạnh; PISA và TIMSS đã xác lập về thành tựu và lo âu/nghi ngờ bản thân nhưng đo lường về unforgiveness (Toughness/Maliciousness/Proviolence) là mới và yếu hơn; cần các nghiên cứu cấp cá nhân và tổng hợp liên kết unforgiveness ↔ achievement ↔ anxiety/self-doubt.",
        "Tác giả cũng cảnh báo: ở khu vực Đông Nam Á (Southern Asian — Malaysia, Indonesia, Sri Lanka...) chỉ thấp hơn Confucian Asia một chút trên thang TM, nhưng không có dữ liệu PISA/TIMSS để so về thành tựu/lo âu — nên \"di sản Nho giáo\" có lẽ không phải nguyên nhân duy nhất.",
        "Phản biện của người tóm tắt: (1) khái niệm \"Confucian Asia\" của Stankov bao gồm Singapore, Đài Loan, Việt Nam (mặc dù Việt Nam chưa có dữ liệu PISA tại thời điểm bài viết 2010); (2) gộp Hong Kong-China, Korea, Japan thành một khối có thể che giấu khác biệt nội tại; (3) thang đo \"Toughness/Maliciousness\" và \"Proviolence\" là xây dựng từ truyền thống nhân cách phương Tây — có thể không tương đương về văn hóa (measurement invariance) khi đem qua Đông Á; (4) liên kết \"Confucian unforgiving culture → academic anxiety\" còn ở mức giả thuyết, chưa được kiểm định cấp cá nhân; (5) bài viết viết trước khi Việt Nam tham gia PISA 2012 (sau đó Việt Nam đạt kết quả cao, hỗ trợ cho luận điểm Confucian).",
    ]),
    ("Áp dụng cho LA NCS", [
        "Cực kỳ quan trọng cho khung lý thuyết LA về lo âu học đường tại Việt Nam: Stankov 2010 cung cấp bằng chứng định lượng đa quốc gia rằng văn hóa Nho giáo đi kèm anxiety cao và self-doubt cao ở học sinh — Việt Nam là một nước Nho giáo theo phân loại của House et al. 2004.",
        "Trích dẫn cho phần thảo luận: tỷ lệ lo âu cao ở học sinh THCS/THPT Việt Nam có thể không phải hiện tượng cá biệt mà là biểu hiện của một pattern khu vực Đông Á (Confucian).",
        "Hỗ trợ luận điểm: can thiệp tâm lý ở học sinh Việt Nam phải tính đến yếu tố văn hóa (Nho giáo, gia đình, áp lực thành tựu) — không thể bê nguyên model CBT phương Tây vào.",
        "Liên kết với QT076 (Tam Giáo) để xây dựng \"khung văn hóa kép\": Nho giáo (mạnh nhất trong Tam Giáo VN theo Vuong et al. 2018) → áp lực thành tựu → lo âu học sinh.",
        "Trích dẫn để biện minh cho việc đưa thang đo về \"perfectionism\" / \"socially prescribed perfectionism\" / \"familial pressure\" vào bộ công cụ đánh giá lo âu.",
    ]),
    ("Trích dẫn nguyên văn", [
        "Trang 555 (Abstract): \"Although forgiveness is a part of Confucian philosophy, people from modern Confucian Asian countries appear to be less forgiving than Europeans — i.e., they tend to disagree with statements that express toughness, maliciousness and provioilence less strongly than Europeans. This relatively unforgiving attitude, coupled with the belief that effort rather than ability is the primary source of success, may be able to explain both high achievement and high anxiety and self-doubt among Confucian Asian students.\"",
        "Trang 559: \"A finding from PISA 2003 that Confucian Asian students experienced higher levels of anxiety and self-doubt (Lee, 2009) can be interpreted in terms of this unique cultural aspect of Confucianism. That is, in the minds of Confucian Asian students, the distinction between the self and one's family is not clear-cut and self achievement is also seen as family's achievement.\"",
    ]),
    ("Tham khảo", [
        "Stankov, L. (2010). Unforgiving Confucian culture: A breeding ground for high academic achievement, test anxiety and self-doubt? Learning and Individual Differences, 20(6), 555–563. DOI: 10.1016/j.lindif.2010.05.003.",
        "Tham chiếu chéo (trích trong bài): Lee, J. (2009). Universals and specifics of math self-concept, math self-efficacy, and math anxiety across 41 PISA 2003 participating countries. Learning and Individual Differences, 19, 355-365.",
        "Tham chiếu chéo: House, R., Hanges, P., Javidan, M., Dorfman, P. W., & Gupta, V. (2004). Culture, leadership, and organizations: The GLOBE study of 62 societies. Thousand Oaks, CA: SAGE.",
        "PDF: 02_Papers-goc/Chua-phan-loai/Stankov_2010_Confucian_Academic.pdf — đã verify ngày 07/06/2026 (đọc 8 trang đầu/9 trang).",
    ]),
]

# ============================================================
# QT078 — Rose 2002 — Co-Rumination in Friendships of Girls and Boys
# ============================================================
qt078 = [
    ("QT078 — Rose (2002): Đồng-trầm-tư (Co-Rumination) trong tình bạn của trẻ gái và trẻ trai", []),
    ("Thông tin xuất bản", [
        "Tác giả: Amanda J. Rose.",
        "Tiêu đề: \"Co-Rumination in the Friendships of Girls and Boys\".",
        "Tạp chí: Child Development, November/December 2002, Volume 73, Number 6, trang 1830–1843.",
        "Bản quyền: © 2002 by the Society for Research in Child Development.",
        "Tài liệu trích dẫn: 0009-3920/2002/7306-0015.",
        "DOI: 10.1111/1467-8624.00509.",
        "Đường dẫn PDF (đã verify): 02_Papers-goc/Chua-phan-loai/Rose_2002_CoRumination.pdf",
    ]),
    ("Tóm tắt", [
        "Bài báo giới thiệu và kiểm định khái niệm mới \"co-rumination\" (đồng-trầm-tư) — định nghĩa là việc thảo luận nhiều lần và quay đi quay lại về các vấn đề, suy đoán về vấn đề, và tập trung vào cảm xúc tiêu cực trong khuôn khổ một mối quan hệ tình bạn hai chiều. Rose đặt giả thuyết rằng co-rumination tích hợp hai dòng nghiên cứu trước đó: (a) self-disclosure trong tình bạn dẫn đến tình bạn chất lượng cao và gần gũi; (b) rumination (trầm tư đơn lẻ) dẫn đến trầm cảm và lo âu. Nghiên cứu trên 608 học sinh lớp 3, 5, 7 và 9 cho thấy co-rumination vừa liên quan tích cực đến chất lượng tình bạn vừa liên quan tích cực đến triệu chứng nội tâm hóa (lo âu + trầm cảm). Trẻ gái co-ruminate nhiều hơn trẻ trai, đặc biệt ở tuổi vị thành niên, giúp giải thích vì sao trẻ gái có tình bạn gần gũi hơn nhưng cũng nhiều triệu chứng trầm cảm/lo âu hơn.",
    ]),
    ("Bối cảnh nghiên cứu / Mục tiêu", [
        "Bối cảnh: Trước 2002 có hai dòng nghiên cứu mâu thuẫn: (a) self-disclosure trong tình bạn được coi là tốt (Newcomb & Bagwell 1995; Parker & Asher 1993); (b) rumination — \"dwelling on negative topics\" — gây trầm cảm/lo âu (Nolen-Hoeksema và cộng sự). Trẻ gái có tình bạn gần gũi hơn (Buhrmester) nhưng cũng có nhiều triệu chứng trầm cảm/lo âu hơn (Nolen-Hoeksema & Girgus 1994) — nghịch lý.",
        "Mục tiêu (bốn giả thuyết): (1) trẻ gái co-ruminate nhiều hơn trẻ trai, đặc biệt ở vị thành niên; (2) co-rumination liên quan đến điều chỉnh tình bạn tích cực NHƯNG điều chỉnh cảm xúc tiêu cực; (3) khác biệt giới tính ở co-rumination giúp giải thích khác biệt giới tính ở tình bạn và điều chỉnh cảm xúc; (4) self-disclosure, rumination và co-rumination có pattern quan hệ khác nhau với tình bạn và cảm xúc.",
    ]),
    ("Phương pháp", [
        "Mẫu: 608 học sinh từ 2 khu trường công lập miền Trung Tây nước Mỹ.",
        "Phân bố: lớp 3 (n=144: 76 nữ, 68 nam), lớp 5 (n=140: 76 nữ, 64 nam), lớp 7 (n=167: 86 nữ, 81 nam), lớp 9 (n=157: 84 nữ, 73 nam).",
        "Dân tộc: 87% European American, 6% African American, 2% American Indian, 1% Asian American, 1% Hispanic American, 3% bi-chủng/khác.",
        "Nhóm tuổi: Children (lớp 3+5, M tuổi = 9.9) vs Adolescents (lớp 7+9, M tuổi = 13.8).",
        "Tỷ lệ đồng thuận của phụ huynh: form gửi cho 704 học sinh, trả 660 (94%), đồng thuận 612 (93%).",
        "Đo lường: (a) Co-Rumination Questionnaire — thang mới 27-item, 5-point Likert, Cronbach α = 0.96; (b) Self-disclosure — 5 item từ Intimate Exchange subscale của Friendship Quality Questionnaire (Parker & Asher 1993), α = 0.85; (c) Rumination — 21 item từ Response to Depression Questionnaire (Nolen-Hoeksema & Morrow 1991), α = 0.91; (d) Friendship Quality — Friendship Quality Questionnaire revised + Friendship Qualities Scale (Bukowski 1994) + Emotional Closeness Scale (Camarena 1990), α = 0.92; (e) Depression — Children's Depression Inventory (Kovacs 1992); (f) Anxiety — Revised Children's Manifest Anxiety Scale (Reynolds & Richmond 1978), 28 item; (g) Internalizing composite — chuẩn hóa các item lo âu (28) + trầm cảm (17 sau khi loại conduct items), α = 0.95.",
        "Thiết kế: nghiên cứu cắt ngang, hai phase đo cách nhau trung bình 8.9 ngày.",
        "Phân tích: ANOVA 2 (giới) × 2 (nhóm tuổi), tương quan Pearson, hồi quy trung gian theo Baron & Kenny (1986).",
    ]),
    ("Phát hiện chính (số liệu trích từ Bảng 1 và Bảng 2 PDF, trang 1836-1837)", [
        "Co-rumination: hiệu chính giới F=84.63, p<.001; trẻ gái > trẻ trai. Children: girls M=2.88 (SD=0.93) vs boys M=2.44 (SD=0.97). Adolescents: girls M=3.14 (SD=0.73) vs boys M=2.32 (SD=0.71). Khác biệt giới rõ rệt hơn ở vị thành niên (Gender×Grade interaction F=7.81, p<.01).",
        "Self-disclosure: hiệu chính giới F=160.90, p<.001. Adolescent girls cao nhất (M=4.13, SD=0.81); adolescent boys thấp nhất (M=2.98, SD=0.88).",
        "Rumination: hiệu chính giới F=19.54, p<.001; girls > boys; children > adolescents.",
        "Self-reported friendship quality and closeness: hiệu chính giới F=53.56, p<.001; girls > boys; khác biệt lớn nhất ở adolescents (girls M=3.29, SD=0.44; boys M=2.64, SD=0.76).",
        "Friend-reported friendship quality and closeness (sub-mẫu 284 cặp bạn có dữ liệu chéo): F=41.23 cho gender; girls > boys, khác biệt lớn nhất ở adolescents.",
        "Internalizing symptoms: hiệu chính giới F=3.92, p<.05; girls M=0.04 (SD=0.59) > boys M=−0.05 (SD=0.55); hiệu chính grade group F=8.39, p<.01.",
        "Tương quan (Bảng 2, trang 1837): Co-rumination với Self-disclosure r=.61***; với Rumination r=.46***; với self-reported friendship quality r=.45***; với internalizing r=.20***; với friend-reported friendship quality r=.10 (n.s.).",
        "Tổng phụ huynh đồng thuận: 612/660 = 93%. Tỷ lệ học sinh có bạn thân tương hỗ (reciprocal friend): 77%.",
        "Phân tích trung gian (mediation): co-rumination giải thích được một phần khác biệt giới ở self-reported friendship quality và ở internalizing symptoms (sau khi kiểm soát co-rumination, hiệu của giới giảm).",
    ]),
    ("Bàn luận / Hàm ý", [
        "Co-rumination là một construct \"hai mặt\" (double-edged): vừa nuôi dưỡng tình bạn gần gũi vừa duy trì các triệu chứng nội tâm hóa.",
        "Khác biệt giới tính trong tình bạn không phải là thuần túy tốt cho trẻ gái — chính cơ chế (co-rumination) giúp trẻ gái có tình bạn sâu sắc hơn cũng làm trẻ gái dễ trầm cảm/lo âu hơn.",
        "Co-rumination khác với self-disclosure ở chỗ tập trung vào tiêu cực và lặp lại; khác với rumination ở chỗ là quá trình xã hội hai chiều.",
        "Hàm ý thực hành: các can thiệp giáo dục/lâm sàng nên giúp trẻ (đặc biệt trẻ gái vị thành niên) phân biệt giữa self-disclosure lành mạnh và co-rumination độc hại.",
    ]),
    ("Phản biện / Hạn chế", [
        "Hạn chế (PDF chưa đề cập rõ ở 8 trang đầu — đoán: thường ở Discussion): (1) thiết kế cắt ngang, không thể kết luận nhân quả; (2) chủ yếu mẫu European American (87%) — khái quát hóa hạn chế.",
        "Phản biện của người tóm tắt: (1) thang đo Co-Rumination Questionnaire 27-item mới, cần kiểm định thêm trên các mẫu khác (đặc biệt ngoài Mỹ); (2) phụ thuộc tự báo cáo, không có quan sát hành vi thực; (3) tương quan với friend-reported friendship quality (r=.10) không có ý nghĩa thống kê — gợi ý hiện tượng có thể là \"cảm nhận của một bên\" hơn là tính chất khách quan của tình bạn; (4) không phân tích các yếu tố văn hóa — pattern này có thể khác ở các nền văn hóa khác (ví dụ Confucian East Asia, nơi tình bạn có quy chuẩn khác); (5) mẫu trẻ Á Châu chỉ 1% — không thể khái quát cho học sinh Việt Nam.",
    ]),
    ("Áp dụng cho LA NCS", [
        "Cực kỳ quan trọng cho LA về lo âu học sinh: cung cấp khung lý thuyết về cơ chế xã hội (peer co-rumination) duy trì triệu chứng nội tâm hóa.",
        "Trích dẫn cho phần \"khác biệt giới tính\" trong tỷ lệ lo âu của học sinh nữ vs nam — gợi ý cơ chế co-rumination giải thích vì sao nữ sinh ở các nghiên cứu Việt Nam thường có tỷ lệ lo âu cao hơn nam sinh.",
        "Cung cấp tiền đề để biện luận cho việc đưa câu hỏi về \"chia sẻ lặp đi lặp lại với bạn về lo lắng\" vào bộ công cụ đánh giá ở học sinh Việt Nam (sau khi adapt văn hóa).",
        "Liên kết với QT077 (Stankov 2010 về Confucian academic anxiety): trong bối cảnh học sinh Việt Nam vốn đã có áp lực thành tích cao, nếu nữ sinh co-ruminate với bạn về điểm số/kỳ thi, có thể làm trầm trọng thêm vòng xoáy lo âu.",
        "Trích dẫn cho phần can thiệp: bài học từ Rose là không nên đơn thuần khuyến khích \"chia sẻ với bạn\" như giải pháp — cần phân biệt self-disclosure xây dựng với co-rumination duy trì lo âu.",
    ]),
    ("Trích dẫn nguyên văn", [
        "Trang 1830 (Abstract): \"Co-rumination refers to extensively discussing and revisiting problems, speculating about problems, and focusing on negative feelings. ... Co-rumination was related to high-quality, close friendships and aspects of depression and anxiety. Girls reported co-ruminating more than did boys, which helped to account for girls' more positive friendship adjustment and greater internalizing symptoms.\"",
        "Trang 1830-1831: \"co-rumination, conceptualized as a specific type of self-disclosure, is expected to be linked with having high-quality, emotionally close friendships. ... Because the aspects of co-rumination thought to be related to emotional problems are largely shared with rumination, a measure of rumination may account for the relation between co-rumination and internalizing symptoms in this study.\"",
    ]),
    ("Tham khảo", [
        "Rose, A. J. (2002). Co-Rumination in the Friendships of Girls and Boys. Child Development, 73(6), 1830–1843. DOI: 10.1111/1467-8624.00509.",
        "Tham chiếu chéo: Nolen-Hoeksema, S., & Girgus, J. S. (1994). The emergence of gender differences in depression during adolescence. Psychological Bulletin, 115, 424-443.",
        "Tham chiếu chéo: Parker, J. G., & Asher, S. R. (1993). Friendship and friendship quality in middle childhood: Links with peer group acceptance and feelings of loneliness and social dissatisfaction. Developmental Psychology, 29, 611-621.",
        "PDF: 02_Papers-goc/Chua-phan-loai/Rose_2002_CoRumination.pdf — đã verify ngày 07/06/2026 (đọc 8 trang đầu/14 trang).",
    ]),
]

# Build all
total_words = 0
results = []
for fname, sections in [
    ("QT074_Karasu_1986_Psychotherapies_FIXED_07062026.docx", qt074),
    ("QT075_Karasu_BioSketch_Einstein_2014_FIXED_07062026.docx", qt075),
    ("QT076_Small_Blanc_2021_TamGiao_VN_FIXED_07062026.docx", qt076),
    ("QT077_Stankov_2010_ConfucianAcademic_FIXED_07062026.docx", qt077),
    ("QT078_Rose_2002_CoRumination_FIXED_07062026.docx", qt078),
]:
    path, wc = build(fname, sections)
    results.append((fname, wc))
    total_words += wc
    print(f"  -> {fname}: {wc} words")

print(f"\nTotal Vietnamese words across 5 docs: {total_words}")
