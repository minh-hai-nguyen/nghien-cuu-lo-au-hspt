# -*- coding: utf-8 -*-
"""Index 4 doc mới (V1 + Herres + Steinhoff + I²) vào RAG + KG + verify."""
import sys, io, os, json, re
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document

KG_DIR  = r'c:\Users\OS\OneDrive\read_books\Lo-au\06_Scripts\kg_data'
RAG_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\rag_dich_phan_bien'
COLLECTION = 'lo_au_dich_phan_bien'

DOCS = [
    {
        'file': r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\Tra_loi_HieuQua_CopingTuPhat_HS_cho_thay_27042026.docx',
        'doc_id': 'TraLoi_HieuQuaCoping_v1', 'topic': 'effectiveness_spontaneous_coping',
        'type_doc': 'qa_for_thay', 'date': '2026-04-27',
    },
    {
        'file': r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\Herres_Ohannessian_2015_CopingProfiles_dich_phan_bien_27042026.docx',
        'doc_id': 'TLN_Herres2015_CopingProfiles', 'topic': 'coping_profiles_LPA',
        'type_doc': 'translation_critique', 'date': '2026-04-27',
        'doi': '10.1016/j.jad.2015.07.031', 'pmid': '26275359', 'pmc': 'PMC4565746',
        'n': 982, 'study_type': 'cross_sectional_LPA',
    },
    {
        'file': r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\Steinhoff_2023_LongitudinalCoping_dich_phan_bien_27042026.docx',
        'doc_id': 'TLN_Steinhoff2023_LongCoping', 'topic': 'longitudinal_coping_predictors',
        'type_doc': 'translation_critique', 'date': '2026-04-27',
        'doi': '10.1177/02724316231181660', 'pmid': '39372429', 'pmc': 'PMC10261967',
        'n': 786, 'study_type': 'longitudinal_cohort_9y',
    },
    {
        'file': r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\I_binh_phuong_giai_thich_cho_thay_27042026.docx',
        'doc_id': 'GiaiThich_I2', 'topic': 'I_squared_heterogeneity_meta_analysis',
        'type_doc': 'methodology_explanation', 'date': '2026-04-27',
    },
]

# === BƯỚC 1 — CHUNKING ===
print('=== BƯỚC 1: CHUNKING ===')
def docx_to_chunks(path):
    d = Document(path)
    chunks = []; current_h = ''; current_text = []
    for p in d.paragraphs:
        is_heading = False
        for r in p.runs:
            if r.bold and r.font.size and r.font.size.pt >= 13:
                is_heading = True; break
        text = p.text.strip()
        if not text: continue
        if is_heading and (text.startswith('PHẦN') or text.startswith('CÂU') or
                           re.match(r'^\d+\.', text) or text.startswith('Bài') or
                           text.startswith('H') or text.isupper()):
            if current_text: chunks.append((current_h, '\n'.join(current_text)))
            current_h = text; current_text = [text]
        else:
            current_text.append(text)
    if current_text: chunks.append((current_h, '\n'.join(current_text)))
    for ti, t in enumerate(d.tables):
        rows_txt = []
        for row in t.rows:
            rows_txt.append(' | '.join(c.text.strip() for c in row.cells))
        chunks.append((f'BẢNG {ti+1}', '\n'.join(rows_txt)))
    return [(h, t) for h, t in chunks if len(t) > 200]

all_chunks = []
for doc_meta in DOCS:
    chunks = docx_to_chunks(doc_meta['file'])
    print(f"  {doc_meta['doc_id']}: {len(chunks)} chunks")
    for hi, (h, t) in enumerate(chunks):
        all_chunks.append({
            'chunk_id': f"{doc_meta['doc_id']}_chunk_{hi:02d}",
            'doc_id': doc_meta['doc_id'], 'section': h[:80],
            'text': t[:3500], 'metadata': doc_meta,
        })
print(f'  TỔNG: {len(all_chunks)} chunks')

# === BƯỚC 2 — RAG ===
print()
print('=== BƯỚC 2: ADD TO RAG ===')
os.environ['HF_HOME'] = r'D:\hf_cache'
import chromadb
from sentence_transformers import SentenceTransformer
model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
client = chromadb.PersistentClient(path=RAG_DIR)
col = client.get_collection(COLLECTION)
print(f'  Trước: {col.count()} chunks')
ids = [c['chunk_id'] for c in all_chunks]
docs_text = []; metas = []
for c in all_chunks:
    md = c['metadata']
    full_text = f"DOC {c['doc_id']} — Topic: {md['topic']} — Type: {md['type_doc']}\n--- {c['section']} ---\n{c['text']}"
    docs_text.append(full_text)
    meta = {'doc_id': c['doc_id'], 'section': c['section'], 'topic': md['topic'],
            'type_doc': md['type_doc'], 'date': md['date'], 'source': 'coping_27042026'}
    if 'doi' in md: meta['doi'] = md['doi']
    if 'pmid' in md: meta['pmid'] = md['pmid']
    if 'pmc' in md: meta['pmc'] = md['pmc']
    metas.append(meta)
emb = model.encode(docs_text, show_progress_bar=False).tolist()
existing = col.get(ids=ids)
if existing['ids']:
    col.delete(ids=existing['ids'])
    print(f'  Deleted {len(existing["ids"])} old (rerun)')
col.add(ids=ids, embeddings=emb, documents=docs_text, metadatas=metas)
print(f'  Sau: {col.count()} chunks (+{len(ids)})')

# === BƯỚC 3 — KG ===
print()
print('=== BƯỚC 3: UPDATE KG ===')
nodes = json.load(open(os.path.join(KG_DIR, 'nodes.json'), encoding='utf-8'))
edges = json.load(open(os.path.join(KG_DIR, 'edges.json'), encoding='utf-8'))
print(f'  Trước: {len(nodes)} nodes, {len(edges)} edges')
new_nodes = []; new_edges = []
for md in DOCS:
    did = md['doc_id']
    if any(n.get('id') == did for n in nodes): continue
    node = {'id': did, 'type': 'Document', 'topic': md['topic'],
            'doc_type': md['type_doc'], 'date': md['date']}
    if 'doi' in md: node['doi'] = md['doi']
    if 'pmid' in md: node['pmid'] = md['pmid']
    if 'pmc' in md: node['pmc'] = md['pmc']
    if 'n' in md: node['sample_size'] = md['n']
    if 'study_type' in md: node['study_type'] = md['study_type']
    new_nodes.append(node)
    new_edges.append({'source': did, 'target': f"Topic::{md['topic']}", 'rel': 'ABOUT_TOPIC'})
    if 'doi' in md: new_edges.append({'source': did, 'target': f"DOI::{md['doi']}", 'rel': 'HAS_DOI'})
    if 'pmid' in md: new_edges.append({'source': did, 'target': f"PMID::{md['pmid']}", 'rel': 'HAS_PMID'})
    if 'pmc' in md: new_edges.append({'source': did, 'target': f"PMC::{md['pmc']}", 'rel': 'HAS_PMC'})
    if 'n' in md: new_edges.append({'source': did, 'target': f"SampleSize::{md['n']}", 'rel': 'HAS_N'})

# Cross-references
new_edges.extend([
    {'source': 'TLN_Herres2015_CopingProfiles', 'target': 'QT08', 'rel': 'SHARES_METHOD_LPA'},
    {'source': 'TLN_Herres2015_CopingProfiles', 'target': 'TraLoi_HieuQuaCoping_v1', 'rel': 'CITED_IN'},
    {'source': 'TLN_Steinhoff2023_LongCoping', 'target': 'TraLoi_HieuQuaCoping_v1', 'rel': 'CITED_IN'},
    {'source': 'TLN_Herres2015_CopingProfiles', 'target': 'Concept::adaptive_coping', 'rel': 'STUDIES_CONCEPT'},
    {'source': 'TLN_Herres2015_CopingProfiles', 'target': 'Concept::maladaptive_coping', 'rel': 'STUDIES_CONCEPT'},
    {'source': 'TLN_Herres2015_CopingProfiles', 'target': 'Concept::coping_profile_LPA', 'rel': 'INTRODUCES_CONCEPT'},
    {'source': 'TLN_Steinhoff2023_LongCoping', 'target': 'Concept::inoculation_hypothesis', 'rel': 'TESTS_CONCEPT'},
    {'source': 'TLN_Steinhoff2023_LongCoping', 'target': 'Concept::supportive_parenting', 'rel': 'STUDIES_CONCEPT'},
    {'source': 'GiaiThich_I2', 'target': 'Concept::heterogeneity', 'rel': 'EXPLAINS_CONCEPT'},
    {'source': 'GiaiThich_I2', 'target': 'Concept::meta_analysis', 'rel': 'EXPLAINS_CONCEPT'},
    {'source': 'TraLoi_HieuQuaCoping_v1', 'target': 'TLN_Herres2015_CopingProfiles', 'rel': 'SUMMARIZES'},
    {'source': 'TraLoi_HieuQuaCoping_v1', 'target': 'TLN_Steinhoff2023_LongCoping', 'rel': 'SUMMARIZES'},
])

nodes.extend(new_nodes); edges.extend(new_edges)
json.dump(nodes, open(os.path.join(KG_DIR, 'nodes.json'), 'w', encoding='utf-8'),
          ensure_ascii=False, indent=2)
json.dump(edges, open(os.path.join(KG_DIR, 'edges.json'), 'w', encoding='utf-8'),
          ensure_ascii=False, indent=2)
print(f'  Sau: {len(nodes)} nodes (+{len(new_nodes)}), {len(edges)} edges (+{len(new_edges)})')

# === BƯỚC 4 — VERIFY ===
print()
print('=== BƯỚC 4: VERIFY (RAG queries với metadata filter) ===')
queries = [
    ('hiệu quả coping tự phát học sinh adaptive maladaptive', 'TraLoi_HieuQuaCoping_v1'),
    ('Herres Ohannessian 2015 4 hồ sơ coping LPA', 'TLN_Herres2015_CopingProfiles'),
    ('Steinhoff 2023 longitudinal cohort 9 năm Zurich z-proso', 'TLN_Steinhoff2023_LongCoping'),
    ('I bình phương heterogeneity meta-analysis Higgins Thompson', 'GiaiThich_I2'),
    ('inoculation hypothesis miễn dịch tâm lý', 'TLN_Steinhoff2023_LongCoping'),
    ('Disengaged Independent Active Social Support Seeking 982 học sinh', 'TLN_Herres2015_CopingProfiles'),
    ('I² < 25% thấp ngưỡng MA random fixed effects', 'GiaiThich_I2'),
]
correct = 0
for q, expected in queries:
    eq = model.encode([q]).tolist()
    r = col.query(query_embeddings=eq, n_results=2, where={'doc_id': expected},
                  include=['metadatas','distances'])
    if r['metadatas'] and r['metadatas'][0]:
        top = r['metadatas'][0][0]
        ok = '✓' if top['doc_id'] == expected else ' '
        if ok == '✓': correct += 1
        sec = top['section'][:55]
        did = top['doc_id']
        qprev = q[:50]
        print(f'  {ok} "{qprev}" → {did} | {sec}')
print(f'  Top-1 với filter: {correct}/{len(queries)} ({100*correct/len(queries):.0f}%)')

# KG check
print()
print('=== KG QUERIES ===')
for n in [n for n in nodes if n.get('id', '').startswith(('TraLoi_HieuQua','TLN_Herres','TLN_Steinhoff','GiaiThich_I2'))]:
    edges_from = [e for e in edges if e['source']==n['id']]
    edges_to = [e for e in edges if e['target']==n['id']]
    nid = n['id']
    print(f'  {nid}: {len(edges_from)} outgoing + {len(edges_to)} incoming')

print()
print('HOÀN TẤT.')
