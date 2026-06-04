# -*- coding: utf-8 -*-
"""Verify LA v3_2 - check loi da sua."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_2_FixSoLieu_27052026.docx')
d = Document(LA)

print(f"=== Verify LA v3_2 ===")
print(f"Paragraphs: {len(d.paragraphs)}, Tables: {len(d.tables)}")

# Sanity checks
checks = [
    ('vùng bán đô thị', 'Lâm context - should be 0 in body', 0),
    ('bán nông thôn', 'Lâm context - should be >= 2', None),
    ('rối loạn lo âu tổng quát', 'should be 0', 0),
    ('rối loạn lo âu lan tỏa', 'should be high', None),
    ('lo âu tổng quát', 'should be 0', 0),
    ('lo âu lan tỏa', 'should be high', None),
    ('RLLATQ', 'should be 0', 0),
    ('RLLALT', 'should be 23+', None),
    ('Lê Minh T., Nguyễn Đăng K., Ngô Anh V.', 'should be 0', 0),
    ('Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh', 'should be 1', None),
    ('có 49% học sinh lo âu', 'should be 0', 0),
    ('có 49,0% học sinh lo âu', 'should be 1', None),
    ('7,7% loại nhẹ, 24,5% loại vừa', 'should be 0', 0),
    ('11,2% loại nhẹ, 25,1% loại vừa', 'should be 1', None),
]

for kw, desc, exp in checks:
    # Count in body paragraphs
    body_count = sum(1 for p in d.paragraphs if kw in p.text)
    # Count in table cells
    tbl_count = 0
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                if kw in cell.text:
                    tbl_count += 1
    total = body_count + tbl_count
    flag = '✓' if (exp is None) or (total == exp) else '✗'
    print(f"  {flag} '{kw}' body={body_count} tables={tbl_count} total={total} (expect {exp}) - {desc}")

# Also verify the para 271 (An Giang) and para 1368 (TLTK)
print(f"\n--- para 271 ---")
print(d.paragraphs[271].text[:300])
print(f"\n--- para 1368 ---")
print(d.paragraphs[1368].text[:300])
print(f"\n--- para 267 ---")
print(d.paragraphs[267].text[:400])
print(f"\n--- para 268 ---")
print(d.paragraphs[268].text[:400])
print(f"\n--- para 269 ---")
print(d.paragraphs[269].text)
