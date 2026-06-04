# -*- coding: utf-8 -*-
"""
Dịch đầy đủ VN21 — Trần Thảo Vi et al. 2024
Academic stress among students in Vietnam: a three-year longitudinal study
Journal of Rural Medicine 19(4): 279-290
"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()

# ========== 1. LINK ==========
add_p(doc, 'Link bài báo gốc: https://doi.org/10.2185/jrm.2024-012', size=10)
add_p(doc, 'PMC Open Access: https://pmc.ncbi.nlm.nih.gov/articles/PMC11442093/', size=10)

# ========== 2. TIÊU ĐỀ ==========
add_heading(doc, 'Căng thẳng học tập ở học sinh tại Việt Nam: Nghiên cứu dọc ba năm về tác động của các yếu tố gia đình, lối sống và học tập', 1)
h = doc.add_paragraph()
r = h.add_run('Academic stress among students in Vietnam: a three-year longitudinal study on the impact of family, lifestyle, and academic factors')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

# ========== 3. THÔNG TIN THƯ MỤC ==========
add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Academic stress among students in Vietnam: a three-year longitudinal study on the impact of family, lifestyle, and academic factors'),
    ('Tiêu đề dịch', 'Căng thẳng học tập ở học sinh tại Việt Nam: Nghiên cứu dọc ba năm về tác động của các yếu tố gia đình, lối sống và học tập'),
    ('Tác giả', 'Thao Vi Tran (1,2), Hoang Thuy Linh Nguyen (3), Xuan Minh Tri Tran (1,3), Yuri Tashiro (1), Kaoruko Seino (1), Thang Van Vo (3), Keiko Nakamura (1)'),
    ('Cơ quan', '(1) Khoa Khởi nghiệp Y tế Toàn cầu, ĐH Y Nha khoa Tokyo (TMDU), Nhật Bản\n(2) Viện Nghiên cứu Sức khỏe Cộng đồng, ĐH Y Dược Huế, ĐH Huế, Việt Nam\n(3) Khoa Y tế Công cộng, ĐH Y Dược Huế, ĐH Huế, Việt Nam'),
    ('Tạp chí', 'Journal of Rural Medicine (J Rural Med) — Hiệp hội Y học Nông thôn Nhật Bản'),
    ('Thông tin xuất bản', '2024, Vol. 19(4), pp. 279–290, 12 trang'),
    ('DOI', '10.2185/jrm.2024-012'),
    ('Nhận bài / Chấp nhận', '11/03/2024 / 05/07/2024'),
    ('Loại NC', 'Nghiên cứu thuần tập DỌC (longitudinal cohort) — 3 năm (2018–2021)'),
    ('Mẫu', '611 HS ban đầu → 341 HS phân tích cuối cùng, lớp 6 (2018) → lớp 9 (2021), 4 trường THCS công lập, TP Huế, Việt Nam'),
    ('Tài trợ', 'Quỹ Khoa học Xã hội Nhật Bản (JSPS Grant 17H02164)'),
    ('Truy cập', 'Open Access — CC BY-NC-ND 4.0'),
])

add_page_ref(doc, '279–290', 'Journal of Rural Medicine', 'Vol. 19(4), 2024')

# ========== 4. TÓM TẮT (Abstract) ==========
add_heading(doc, 'TÓM TẮT', 2)

add_p(doc, 'Mục tiêu: Căng thẳng học tập (academic stress) có liên quan đến các rối loạn sức khỏe tâm thần, đặc biệt là trầm cảm và lo âu ở học sinh. Giảm thiểu căng thẳng có thể giảm tỷ lệ rối loạn sức khỏe tâm thần và cải thiện hạnh phúc của học sinh. Nghiên cứu này khám phá các yếu tố ảnh hưởng đến căng thẳng học tập ở học sinh trung học cơ sở tại Việt Nam.')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Một nghiên cứu dọc ba năm được thực hiện sử dụng bảng hỏi tự báo cáo với 611 học sinh từ bốn trường THCS tại thành phố Huế, Việt Nam. Căng thẳng học tập được đánh giá bằng Thang Căng thẳng Giáo dục dành cho Thanh thiếu niên (ESSA — Educational Stress Scale for Adolescents). Các yếu tố gia đình bao gồm số anh chị em và trình độ học vấn của cha mẹ; yếu tố lối sống bao gồm hoạt động thể chất và giấc ngủ; và yếu tố học tập bao gồm điểm trung bình và tham gia học thêm. Mô hình hồi quy tuyến tính được sử dụng để phân tích mối liên quan giữa điểm ESSA tại thời điểm theo dõi và các yếu tố gia đình, lối sống, học tập tại thời điểm ban đầu.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
r = p.add_run('Kết quả: Tổng cộng 341 học sinh hoàn thành cả hai đợt khảo sát ban đầu và theo dõi và trả lời các câu hỏi cần thiết cho phân tích này. Điểm ESSA trung bình của 341 học sinh TĂNG từ 46,4 ± 7,6 (trung bình ± SD) lên 53,5 ± 10,8, từ năm 2018 đến 2021 (p < 0,001). Mô hình đa biến cho thấy: số anh chị em (β = 2,24; KTC 95%: 0,92–3,57; p < 0,01), trình độ học vấn cao của cha (β = 3,20; KTC 95%: 0,13–6,27; p < 0,05), giới tính nữ (β = −2,85 cho nam so với nữ; KTC 95%: −5,15 đến −0,54; p < 0,05), điểm học tập thấp hơn (β = −1,79; KTC 95%: −3,02 đến −0,56; p < 0,01), và tham gia học thêm (β = 4,73; KTC 95%: 0,41–9,06; p < 0,05) có liên quan đến căng thẳng học tập tổng thể. Ngược lại, không có mối liên quan giữa các biến lối sống, sự tập trung của cha mẹ, và sự chấp nhận của cha mẹ với căng thẳng học tập tổng thể.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_p(doc, 'Kết luận: Các phát hiện nhấn mạnh tác động của yếu tố gia đình và khối lượng học tập lên căng thẳng học tập, nhấn mạnh nhu cầu chăm sóc phù hợp từ gia đình và nhà trường để giảm thiểu hoặc ngăn ngừa căng thẳng học tập ở học sinh và cung cấp cho các em một môi trường học tập thoải mái và lành mạnh.')

add_p(doc, 'Từ khóa: căng thẳng học tập, Thang Căng thẳng Giáo dục dành cho Thanh thiếu niên (ESSA), học sinh trung học cơ sở.')

# ========== 5. TÓM TẮT ĐÁNH GIÁ NHANH ==========
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'NC DỌC (longitudinal) 3 năm — HIẾM CỰC KỲ tại VN. Theo dõi cùng nhóm HS từ lớp 6 (11 tuổi) đến lớp 9 (14 tuổi).',
    'Hue Healthy Adolescent Cohort Study — thuần tập trường học tại Huế, hợp tác TMDU Nhật Bản.',
    'ESSA (α = 0,83–0,88) — thang đo căng thẳng giáo dục chuẩn hóa, đã xác thực tại VN (Truc et al. 2015).',
    'Căng thẳng TĂNG đáng kể: ESSA M = 46,4 → 53,5 (tăng 15,3% trong 3 năm, p < 0,001).',
    'Yếu tố mạnh nhất: HỌC THÊM (β = 4,73) — phù hợp bối cảnh VN (HS THCS phải thi vào lớp 10).',
    'Cha học vấn cao → con stress cao hơn (β = 3,20) — phát hiện thú vị, phản ánh kỳ vọng gia đình châu Á.',
    'HS CẤP 2 (THCS) — ít NC tại VN về nhóm tuổi này (đa số NC tập trung cấp 3). Đặc biệt phù hợp đề tài.',
    'Open Access (CC BY-NC-ND 4.0) — tải miễn phí.',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'n = 341 (phân tích cuối) — nhỏ so với NC cắt ngang (Hoa 2024: n=3.910). Nhưng chấp nhận được cho NC dọc.',
    'Follow-up trùng COVID-19 (2021) — có thể ảnh hưởng kết quả (tăng stress do dịch, không chỉ do học tập).',
    'Chỉ TP Huế (đô thị) — không đại diện nông thôn/DTTS. So: Ngô Anh Vinh 2024 (VN15): DTTS Lạng Sơn.',
    'ESSA đo CĂNG THẲNG HỌC TẬP — không đo lo âu trực tiếp (GAD-7, DASS-21). Liên quan nhưng khác biệt.',
    'Tự báo cáo — có thể thiên lệch xã hội. Phỏng vấn mặt đối mặt (20–30 phút) — có thể tốt hơn online nhưng ảnh hưởng riêng tư.',
    'Loại bỏ 176/517 HS (34%) do thiếu dữ liệu — có thể thiên lệch chọn.',
    'Không đo screen time/MXH — yếu tố quan trọng theo Norway 2025 (QT21), Nature 2025 (QT27).',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hướng cải thiện:', bold=True)
for b in [
    'Bổ sung GAD-7/DASS-21 cùng ESSA — đo cả căng thẳng HỌC TẬP và lo âu CHUNG trên cùng mẫu.',
    'Mở rộng: nông thôn + DTTS + TPHCM + Hà Nội — so sánh vùng.',
    'Thêm biến screen time, MXH, giấc ngủ chi tiết (Zhu 2025 QT05: <5h AOR=13,71).',
    'Theo dõi tiếp lớp 10–12 (THPT) — xem căng thẳng tăng hay giảm sau thi vào lớp 10.',
]:
    add_p(doc, f'• {b}')

# ========== 6. GIỚI THIỆU ==========
add_page_ref(doc, '279–280', 'J Rural Med', 'Vol. 19(4), 2024')
add_heading(doc, '1. GIỚI THIỆU', 2)

add_p(doc, 'Gần đây, tác động đáng kể của gánh nặng học tập lên người học ngày càng được thừa nhận, dẫn đến mối quan ngại gia tăng trong xã hội. Một nghiên cứu cho thấy trong các nước OECD, 66% học sinh 15–16 tuổi trải qua căng thẳng về điểm kém và 59% thường xuyên lo lắng về khó khăn khi làm bài kiểm tra (Ribeiro et al., 2018). Nhiều nghiên cứu đã nhấn mạnh mối tương quan giữa căng thẳng học tập và sự xuất hiện các vấn đề sức khỏe tâm thần như trầm cảm, lo âu, và xu hướng tự tử. Căng thẳng quá mức được coi là yếu tố liên quan đến giảm thành tích học tập và tăng tỷ lệ bỏ học (Deng et al., 2022; Pascoe et al., 2020; Jiménez-Mijangos et al., 2023).')

add_p(doc, 'Môi trường trường học là nguồn gây căng thẳng chính cho trẻ em và thanh thiếu niên do áp lực xã hội và học tập, đặc biệt là sự cạnh tranh khốc liệt về thành tích học tập (Klinger et al., 2015). Căng thẳng này được tăng cường bởi nhiều yếu tố, bao gồm khối lượng công việc lớn, lượng tài liệu quá lớn cần nắm bắt, và kỳ vọng về hiệu suất đỉnh cao liên tục. Áp lực học tập được tăng cường thêm bởi kỳ vọng của cha mẹ, nhà trường, và bạn bè; thiếu nguồn lực hỗ trợ thành tích; và ảnh hưởng của yếu tố văn hóa (UNESCO, 2011; Kim et al., 2016).')

add_p(doc, 'Căng thẳng học tập được đặc trưng là trạng thái tâm lý của học sinh phát sinh từ các áp lực xã hội liên tục và tự áp đặt trong môi trường trường học, làm cạn kiệt nguồn lực tâm lý của học sinh (Neseliler et al., 2017). Hơn nữa, căng thẳng học tập được định nghĩa là lo âu và căng thẳng bắt nguồn từ việc học tập và giáo dục (Prabu, 2015).')

add_p(doc, 'Do hậu quả tiềm ẩn của căng thẳng học tập, việc tìm hiểu các yếu tố có thể giảm khả năng thanh thiếu niên đối mặt với căng thẳng học tập là quan trọng trong cả bối cảnh học thuật và lâm sàng. Thành tích học tập đề cập đến mức độ hiểu biết, năng lực và ứng dụng kiến thức trong một môn học cụ thể, thường được đánh giá bởi giáo viên thông qua điểm thi. Một nguồn căng thẳng đáng kể đối với nhiều học sinh bắt nguồn từ áp lực mãnh liệt phải xuất sắc trong các kỳ thi, tạo ra một môi trường học tập căng thẳng cao. Nỗi sợ đạt kết quả dưới kỳ vọng là mối quan ngại phổ biến, khiến một số cá nhân liên kết giá trị bản thân với thành tích học tập (Sansgiry & Sail, 2006; Bedewy & Gabriel, 2015).')

add_p(doc, 'Một số nghiên cứu đã gợi ý rằng anh chị em định hình nhiều khía cạnh phát triển của trẻ em (McHale et al., 2012; Howe & Recchia, 2014; Dirks et al., 2015). Ngoài ra, số lượng anh chị em ảnh hưởng đến căng thẳng học tập của trẻ. Một nghiên cứu tại Trung Quốc cho thấy học sinh con một trải qua mức căng thẳng thấp hơn so với học sinh có anh chị em (Chu et al., 2015). Ngược lại, một nghiên cứu tại Ấn Độ chỉ ra rằng thanh thiếu niên có ít anh chị em hơn phải đối mặt với căng thẳng học tập cao hơn so với những người có nhiều anh chị em hơn (Rentala et al., 2019).')

add_p(doc, 'Kỳ vọng của cha mẹ được xếp hạng trong số các yếu tố gia đình quan trọng liên quan đến căng thẳng giáo dục của học sinh (Talha et al., 2020; Tan & Yates, 2011). Tại Ấn Độ, khoảng 66% học sinh báo cáo rằng trải nghiệm áp lực từ cha mẹ nâng cao thành tích học tập (Deb et al., 2015). Các nghiên cứu khác đã xác nhận rằng kỳ vọng quá cao của cha mẹ có thể khuếch đại căng thẳng học tập ở học sinh (Kim et al., 2016; Ma et al., 2018). Kỳ vọng cao dẫn đến sự can thiệp và kiểm soát quá mức của cha mẹ trong cuộc sống con cái, khiến học sinh phải dành nhiều thời gian hơn cho việc học và đối mặt với nhiều căng thẳng hơn. Mức độ kỳ vọng và áp lực khác biệt đáng kể dựa trên trình độ học vấn của cha mẹ (Deb et al., 2015; Kapali et al., 2019).')

add_p(doc, 'Duy trì chế độ tập luyện thường xuyên và lành mạnh đã được chứng minh mang lại tác động tâm lý tích cực, bao gồm giảm căng thẳng (Monserrat-Hernández et al., 2023). Một nghiên cứu phân tích dữ liệu Mỹ cho thấy những người không tham gia bất kỳ hoạt động thể chất nào trong tuần có khả năng trải qua lo âu và trầm cảm cao hơn gấp đôi so với những người tham gia ít nhất 60 phút tập luyện hàng ngày (Zhu et al., 2019). Dữ liệu từ Khảo sát Quốc gia về Bệnh đồng mắc ở Thanh thiếu niên Mỹ cho thấy mối tương quan đáng kể giữa các kiểu ngủ không đầy đủ và các kết quả tâm lý tiêu cực khác nhau như lo âu, rối loạn hành vi, và giảm sức khỏe tâm thần chủ quan (Zhang et al., 2017).')

add_p(doc, 'Tại Việt Nam, sự chú ý của xã hội đối với thành tích học tập đã tăng lên. Mặc dù có nhiều kỳ vọng khác nhau về thành tích học tập của học sinh, có lo ngại về một số ảnh hưởng tiêu cực của việc theo đuổi thành tích học tập đối với sức khỏe tâm thần của học sinh. Kết quả từ một nghiên cứu trên 1.296 học sinh THCS tại Hà Nội báo cáo rằng tỷ lệ trầm cảm ở học sinh là 27%, với 9,5% và 4,9% trong số đó có trầm cảm nhẹ/trung bình và cực kỳ nặng (Khanh et al., 2023). Một nghiên cứu khác trên 1.161 học sinh THCS tại Việt Nam báo cáo rằng ước tính tỷ lệ lo âu và nguy cơ trầm cảm lần lượt là 22,8% và 41,1% (Nguyen et al., 2013). Trong số những học sinh này, 26,3% đã nghiêm túc cân nhắc tự tử, 12,9% đã xây dựng kế hoạch tự tử, và 3,8% đã cố gắng tự tử.')

add_p(doc, 'Một số nghiên cứu tại Việt Nam đã điều tra mối quan hệ giữa sức khỏe tâm thần và áp lực học tập mà học sinh phải chịu đựng (Nguyen et al., 2013; Thai et al., 2020; Pham et al., 2019). Theo một nghiên cứu năm 2019, học sinh Việt Nam coi điểm số và kỳ thi là có liên quan chặt chẽ đến lo âu (Dinh et al., 2021). Áp lực đạt thành tích học tập tốt hơn có thể được cảm nhận không chỉ bởi học sinh cấp 3 và đại học mà còn bởi học sinh THCS. Điều này là do học sinh THCS phải chuẩn bị cho kỳ thi tốt nghiệp, kết quả kỳ thi này là điều kiện tiên quyết để được tuyển vào trường cấp 3. Điểm đạt được trong kỳ thi này được sử dụng để đánh giá năng lực học tập và xác định cấp độ trường cấp 3 mà học sinh có thể theo học. Điểm cao hơn cho phép vào các trường cấp 3 uy tín hơn, tạo ra một môi trường cạnh tranh làm trầm trọng thêm căng thẳng và áp lực mà học sinh gặp phải trong suốt hành trình giáo dục.')

add_p(doc, 'Mặc dù một số nghiên cứu đã khám phá sức khỏe tâm thần thanh thiếu niên tại Việt Nam, các nghiên cứu về căng thẳng giáo dục và các yếu tố nguy cơ liên quan cho học sinh, đặc biệt là học sinh THCS, vẫn còn khan hiếm. Do đó, nghiên cứu này nhằm điều tra các yếu tố ảnh hưởng đến căng thẳng học tập ở học sinh THCS.')

# ========== 7. PHƯƠNG PHÁP ==========
add_page_ref(doc, '281–283', 'J Rural Med', 'Vol. 19(4), 2024')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)

add_p(doc, '2.1. Người tham gia khảo sát và quy trình', bold=True)
add_p(doc, 'Dữ liệu được thu thập từ Nghiên cứu Thuần tập Thanh thiếu niên Khỏe mạnh Huế (Hue Healthy Adolescent Cohort Study), một nghiên cứu thuần tập ba năm dựa trên trường học (2018–2021) bao gồm học sinh đang theo học tại các trường THCS ở khu vực đô thị Việt Nam. Nghiên cứu thuần tập được phê duyệt bởi Hội đồng Đánh giá Thể chế (IRB) của ĐH Y Nha khoa Tokyo, Nhật Bản, và ĐH Y Dược Huế, Việt Nam. Sở Giáo dục và Đào tạo tỉnh Thừa Thiên Huế, Việt Nam cấp phép tuyển học sinh THCS.')

add_p(doc, 'Khảo sát ban đầu và theo dõi được thực hiện lần lượt vào năm 2018 và 2021. Năm trường THCS công lập từ 23 trường THCS công lập tại TP Huế được chọn sử dụng thiết kế lấy mẫu cụm phân tầng nhiều bậc (multistage, stratified, cluster-random-sampling). Tùy thuộc quy mô mỗi trường, 4–5 lớp học sinh khối 6 (11 tuổi) được chọn ngẫu nhiên. Tất cả học sinh trong các lớp đó có thể giao tiếp hoặc đọc và sẵn lòng tham gia nghiên cứu được tuyển. Học sinh vắng mặt ngày khảo sát bị loại trừ.')

add_p(doc, 'Trong giai đoạn nghiên cứu, đại dịch COVID-19 hoành hành tại TP Huế. Trong năm trường được chọn, bốn trường hoàn thành cả hai đợt khảo sát ban đầu và theo dõi. Người tham gia tại bốn trường được coi là mẫu nghiên cứu cho phân tích này. Tại bốn trường, trước khi thực hiện khảo sát ban đầu, 611 học sinh được mời tham gia. Tại các đợt khảo sát ban đầu và theo dõi năm 2018 và 2021, lần lượt 517 và 479 học sinh tham gia. Sau khi loại trừ dữ liệu không ghi nhận các biến chính, hồ sơ của 341 học sinh được sử dụng trong phân tích. Lý do loại trừ là không có trình độ học vấn của cha hoặc mẹ — các biến quan trọng cho phân tích.')

add_p(doc, 'Nhóm nghiên cứu thu thập dữ liệu thông qua phỏng vấn trực tiếp (face-to-face) sử dụng bảng hỏi có cấu trúc. Mục đích nghiên cứu được giải thích cho học sinh vào ngày thu thập. Học sinh được thông báo dữ liệu sẽ ẩn danh và bảo mật, có thể dừng phỏng vấn bất kỳ lúc nào. Học sinh có sự đồng ý của phụ huynh được phỏng vấn trong phòng khảo sát. Phỏng vấn mất khoảng 20–30 phút.')

add_p(doc, '2.2. Thang đo', bold=True)

add_p(doc, 'Căng thẳng học tập (Academic stress):', bold=True)
add_p(doc, 'Đánh giá bằng Thang Căng thẳng Giáo dục dành cho Thanh thiếu niên (ESSA — Educational Stress Scale for Adolescents), phát triển bởi Sun và cs. (2011). Phiên bản cuối cùng ESSA gồm 16 mục với năm lĩnh vực: "áp lực từ việc học" (pressure from study — phát sinh từ suy nghĩ về giáo dục sau này, nghề nghiệp tương lai, kỳ vọng cha mẹ và nhà trường), "lo lắng về điểm số" (worry about grades — phản ánh sự lo lắng về kết quả kiểm tra), "chán nản" (despondency — thiếu tự tin, không thể tập trung, không hài lòng về thành tích), "tự kỳ vọng" (self-expectations — căng thẳng khi mục tiêu học tập không đạt được), và "khối lượng công việc" (workload — căng thẳng do quá nhiều thời gian học, bài tập, và kiểm tra). Trả lời trên thang Likert 5 điểm (1 = hoàn toàn không đồng ý; 5 = hoàn toàn đồng ý). Tổng điểm ESSA = tổng 16 câu hỏi, dao động 16–80, điểm cao hơn = căng thẳng lớn hơn. Nghiên cứu trước tại VN sử dụng ESSA cho thấy thang đo có tính nhất quán nội bộ tốt (Cronbach α = 0,83; Truc et al., 2015).')

add_p(doc, 'Yếu tố cá nhân — Nhân khẩu gia đình:', bold=True)
add_p(doc, 'Dữ liệu nhân khẩu bao gồm giới tính (nam/nữ), số anh chị em, và trình độ học vấn cha mẹ. Trình độ cha mẹ chia 2 nhóm: thấp (dưới tiểu học, THCS, THPT) và cao (cao đẳng/đại học hoặc sau đại học).')

add_p(doc, 'Yếu tố liên quan học tập:', bold=True)
add_p(doc, 'Điểm trung bình (GPA), tham gia học thêm (extra classes), và gia sư riêng (private tutoring) — có/không.')

add_p(doc, 'Yếu tố lối sống:', bold=True)
add_p(doc, 'Hoạt động thể chất và giờ ngủ đo bằng các biến trong Khảo sát Sức khỏe Học đường Toàn cầu (GSHS). Hoạt động thể chất: bao nhiêu ngày HS hoạt động thể chất ≥60 phút/ngày trong 7 ngày qua (phân loại: <3 ngày và ≥3 ngày). Giờ ngủ trung bình: <8 giờ và ≥8 giờ.')

add_p(doc, 'Thái độ cha mẹ cảm nhận:', bold=True)
add_p(doc, 'Đo bằng 6 câu hỏi liên quan cha mẹ/người giám hộ từ Mô-đun Yếu tố Bảo vệ trong GSHS. Phân tích thành phần chính (varimax rotation) cho 2 chiều: (1) "Sự chấp nhận của cha mẹ" (parental acceptance — cha mẹ coi con là thành viên quan trọng, không quá fixate hay bỏ mặc, khuyến khích con cố gắng hết sức), (2) "Sự tập trung của cha mẹ" (parental concentration — cha mẹ quá bảo vệ, hạn chế khám phá và tương tác xã hội, kỳ vọng con làm nhiều hơn khả năng).')

add_p(doc, '2.3. Phân tích dữ liệu', bold=True)
add_p(doc, 'Thống kê mô tả (%, trung bình, SD) cho ba nhóm: HS tham gia ban đầu, HS tham gia cả hai đợt, HS dùng phân tích cuối. Chênh lệch điểm ESSA (tổng, 5 lĩnh vực, 16 câu hỏi riêng) giữa ban đầu và theo dõi kiểm tra bằng paired sample t-test. Cronbach α đánh giá tính nhất quán nội bộ ESSA. Hồi quy tuyến tính đơn biến và đa biến: mối liên quan giữa căng thẳng học tập tại theo dõi (biến phụ thuộc) và các biến nhân khẩu, học tập, lối sống, thái độ cha mẹ tại ban đầu (biến độc lập). Hệ số beta và KTC 95% được tính. Mô hình đa biến: tất cả biến độc lập được nhập đồng thời. Phân tích thêm mối liên quan giữa các tổ hợp trình độ cha mẹ với căng thẳng. SPSS 25.0, mức ý nghĩa p < 0,05.')

# ========== 8. KẾT QUẢ ==========
add_page_ref(doc, '283–287', 'J Rural Med', 'Vol. 19(4), 2024')
add_heading(doc, '3. KẾT QUẢ', 2)

# Bảng 1: Đặc điểm mẫu
add_heading(doc, 'Bảng 1. Đặc điểm nhân khẩu, học tập, và lối sống của người tham gia', 3)
add_table(doc,
    ['Đặc điểm', 'Ban đầu (n=517) n(%)', 'Theo dõi (n=479) n(%)', 'Phân tích cuối (n=341) n(%)'],
    [['Giới tính: Nam', '269 (52,0%)', '252 (52,6%)', '164 (48,1%)'],
     ['Giới tính: Nữ', '248 (48,0%)', '227 (47,4%)', '177 (51,9%)'],
     ['Học vấn cha — Thấp', '245 (47,4%)', '231 (48,2%)', '190 (55,7%)'],
     ['Học vấn cha — Cao', '189 (36,5%)', '175 (36,5%)', '151 (44,3%)'],
     ['Học vấn mẹ — Thấp', '237 (45,8%)', '225 (47,1%)', '186 (54,5%)'],
     ['Học vấn mẹ — Cao', '199 (38,5%)', '183 (38,2%)', '155 (45,5%)'],
     ['Học thêm — Có', '471 (91,1%)', '438 (91,4%)', '314 (92,1%)'],
     ['Gia sư riêng — Có', '99 (19,1%)', '91 (19,0%)', '64 (18,8%)'],
     ['Thể chất ≥3 ngày', '204 (39,4%)', '191 (39,8%)', '132 (38,7%)'],
     ['Ngủ ≥8 giờ', '453 (87,6%)', '419 (87,5%)', '301 (88,3%)'],
     ['Số anh chị em (M±SD)', '1,51 ± 0,95', '1,49 ± 0,92', '1,43 ± 0,89'],
     ['Điểm TB (M±SD)', '7,40 ± 1,06', '7,43 ± 1,04', '7,61 ± 0,95']],
    widths=[4.0, 3.0, 3.0, 3.5])
add_p(doc, 'Ghi chú: Phân bố biến tương tự giữa 3 nhóm. 92,1% HS học thêm. Chỉ 38,7% hoạt động thể chất ≥3 ngày/tuần.', size=9, italic=True)

# Bảng 2: ESSA scores
add_p(doc, 'Bảng 2 cho thấy trung bình và SD cho tổng điểm ESSA, điểm phụ cho năm lĩnh vực, và điểm cho 16 câu hỏi riêng tại ban đầu và theo dõi ở 341 học sinh.')

add_heading(doc, 'Bảng 2. Điểm ESSA tại ban đầu (2018) và theo dõi (2021) — n = 341', 3)
add_table(doc,
    ['Lĩnh vực / Câu hỏi', 'Ban đầu M (SD)', 'Theo dõi M (SD)', 'p'],
    [['TỔNG ESSA', '46,4 (7,6)', '53,5 (10,8)', '<0,001'],
     ['Áp lực từ việc học (4 câu)', '10,1 (2,8)', '12,7 (3,4)', '<0,001'],
     ['  Q4. Giáo dục/nghề tương lai gây áp lực', '2,5 (1,1)', '3,6 (1,1)', '<0,001'],
     ['  Q5. Cha mẹ quá quan tâm điểm số → áp lực', '2,8 (1,2)', '3,0 (1,2)', '0,029'],
     ['  Q6. Cảm thấy nhiều áp lực trong việc học hàng ngày', '2,5 (1,0)', '3,2 (1,2)', '<0,001'],
     ['  Q11. Quá nhiều cạnh tranh giữa bạn cùng lớp', '2,3 (1,0)', '2,8 (1,1)', '<0,001'],
     ['Lo lắng về điểm số (3 câu)', '11,5 (2,2)', '11,1 (2,6)', '0,008*'],
     ['  Q8. Điểm số rất quan trọng cho tương lai', '4,2 (1,0)', '4,0 (1,1)', '0,060'],
     ['  Q9. Cảm thấy đã làm thất vọng cha mẹ khi điểm kém', '3,8 (1,1)', '3,7 (1,1)', '0,078'],
     ['  Q10. Cảm thấy đã làm thất vọng thầy cô khi điểm kém', '3,5 (1,1)', '3,4 (1,1)', '0,026'],
     ['Chán nản (3 câu)', '8,0 (2,4)', '9,6 (2,8)', '<0,001'],
     ['  Q1. Rất không hài lòng với điểm số', '2,9 (1,3)', '3,2 (1,2)', '<0,001'],
     ['  Q12. Luôn thiếu tự tin về điểm số', '2,8 (1,1)', '3,3 (1,1)', '<0,001'],
     ['  Q13. Rất khó tập trung trong giờ học', '2,3 (0,9)', '3,2 (1,2)', '<0,001'],
     ['Tự kỳ vọng (3 câu)', '8,8 (2,5)', '10,3 (2,9)', '<0,001'],
     ['  Q14. Căng thẳng khi không đạt tiêu chuẩn bản thân', '3,1 (1,1)', '3,7 (1,1)', '<0,001'],
     ['  Q15. Khi không đạt kỳ vọng → cảm thấy không đủ tốt', '3,1 (1,2)', '3,5 (1,2)', '<0,001'],
     ['  Q16. Thường không ngủ được vì lo lắng', '2,6 (1,2)', '3,1 (1,2)', '<0,001'],
     ['Khối lượng công việc (3 câu)', '8,1 (2,4)', '9,8 (2,9)', '<0,001'],
     ['  Q2. Cảm thấy quá nhiều bài vở', '2,8 (1,1)', '3,3 (1,1)', '<0,001'],
     ['  Q3. Cảm thấy quá nhiều bài tập về nhà', '2,3 (0,9)', '3,2 (1,1)', '<0,001'],
     ['  Q7. Cảm thấy quá nhiều bài kiểm tra', '3,0 (1,2)', '3,3 (1,1)', '<0,001']],
    widths=[6.0, 2.5, 2.5, 1.5])
add_p(doc, 'Ghi chú: *Lĩnh vực "lo lắng về điểm số" là lĩnh vực DUY NHẤT GIẢM (11,5 → 11,1). Tất cả lĩnh vực khác TĂNG đáng kể. Cronbach α tổng ESSA = 0,88. Tăng Q13 ("khó tập trung") lớn nhất: 2,3 → 3,2 (+39%).', size=9, italic=True)

# Bảng 3+4: Hồi quy (tóm tắt kết quả chính)
add_p(doc, 'Bảng 3 và 4 cho thấy kết quả phân tích hồi quy tuyến tính đơn biến và đa biến kiểm tra mối quan hệ giữa căng thẳng học tập tại thời điểm theo dõi và các yếu tố liên quan tại thời điểm ban đầu.')

add_heading(doc, 'Bảng 3. Kết quả hồi quy đa biến — Yếu tố liên quan căng thẳng học tập tổng thể', 3)
add_table(doc,
    ['Yếu tố (ban đầu)', 'β', 'KTC 95%', 'p', 'Ý nghĩa'],
    [['Giới tính (nam vs nữ)', '−2,85', '−5,15 đến −0,54', '<0,05', 'Nam THẤP hơn nữ'],
     ['Số anh chị em', '2,24', '0,92–3,57', '<0,01', 'Nhiều ACE → TĂNG stress'],
     ['Học vấn cha (cao vs thấp)', '3,20', '0,13–6,27', '<0,05', 'Cha học cao → con TĂNG stress'],
     ['Học vấn mẹ (cao vs thấp)', '−1,11', '−4,10–1,88', 'n.s.', 'KHÔNG liên quan'],
     ['Điểm học tập', '−1,79', '−3,02 đến −0,56', '<0,01', 'Điểm thấp → TĂNG stress'],
     ['Học thêm (có vs không)', '4,73', '0,41–9,06', '<0,05', 'Học thêm → TĂNG stress (MẠNH NHẤT)'],
     ['Gia sư riêng', '−0,15', '−3,16–2,85', 'n.s.', 'KHÔNG liên quan'],
     ['Thể chất ≥3 ngày', '−0,45', '−2,81–1,92', 'n.s.', 'KHÔNG liên quan'],
     ['Ngủ ≥8 giờ', '−2,49', '−6,05–1,08', 'n.s.', 'KHÔNG liên quan tổng thể'],
     ['Chấp nhận cha mẹ', '0,63', '−0,51–1,76', 'n.s.', 'KHÔNG liên quan'],
     ['Tập trung cha mẹ', '1,12', '−0,04–2,28', 'n.s. (biên)', 'Biên giới — nhưng có ý nghĩa ĐƠN BIẾN']],
    widths=[4.0, 1.5, 3.0, 1.5, 4.0])
add_p(doc, 'Ghi chú: Mô hình đa biến — tất cả biến nhập đồng thời. Học thêm có β lớn nhất (4,73) — HS học thêm có stress tổng thể cao hơn 4,73 điểm ESSA. Giấc ngủ ≥8h có ý nghĩa cho lĩnh vực "chán nản" riêng (β = −1,05, p < 0,05).', size=9, italic=True)

# Bảng 4: Chi tiết theo lĩnh vực
add_heading(doc, 'Bảng 4. Yếu tố có ý nghĩa theo từng lĩnh vực ESSA (đa biến)', 3)
add_table(doc,
    ['Lĩnh vực', 'Yếu tố có ý nghĩa', 'β', 'p'],
    [['Áp lực từ việc học', 'Số ACE', '0,57', '<0,05'],
     ['', 'Điểm học tập', '−0,45', '<0,05'],
     ['', 'Học thêm', '1,68', '<0,05'],
     ['Lo lắng về điểm số', 'Giới tính (nam thấp hơn)', '−0,64', '<0,05'],
     ['', 'Học vấn cha (cao)', '0,74', '<0,05'],
     ['', 'Tập trung cha mẹ', '0,29', '<0,05'],
     ['Chán nản', 'Số ACE', '0,36', '<0,05'],
     ['', 'Học vấn cha (cao)', '0,95', '<0,05'],
     ['', 'Điểm học tập', '−0,99', '<0,001'],
     ['', 'Học thêm', '1,10', '<0,05'],
     ['', 'Ngủ ≥8h', '−1,05', '<0,05'],
     ['Tự kỳ vọng', '(Không có yếu tố có ý nghĩa)', '—', '—'],
     ['Khối lượng công việc', 'Số ACE', '0,81', '<0,001'],
     ['', 'Giới tính (nam thấp hơn)', '−0,76', '<0,05'],
     ['', 'Học vấn cha (cao)', '0,88', '<0,05']],
    widths=[3.5, 4.0, 1.5, 1.5])
add_p(doc, 'Ghi chú: "Tự kỳ vọng" là lĩnh vực DUY NHẤT không có yếu tố nào có ý nghĩa — gợi ý stress tự kỳ vọng mang tính NỘI TẠI, không bị ảnh hưởng bởi yếu tố bên ngoài.', size=9, italic=True)

# Bảng 5: Tổ hợp trình độ cha mẹ
add_heading(doc, 'Bảng 5. Mối liên quan giữa căng thẳng học tập và tổ hợp trình độ cha mẹ', 3)
add_table(doc,
    ['Tổ hợp Cha/Mẹ', 'Đơn biến β (KTC 95%)', 'Đa biến β (KTC 95%)', 'p (đơn biến)'],
    [['Cao/Cao', '0,67 (−1,74; 3,07)', '−4,87 (−10,91; 1,17)', 'n.s.'],
     ['Thấp/Thấp', '−1,77 (−4,08; 0,54)', '−4,87 (−10,91; 1,17)', 'n.s.'],
     ['Cao/Thấp (Cha cao, Mẹ thấp)', '4,43 (0,17; 8,69)', '4,87 (−1,17; 10,91)', '<0,05*'],
     ['Thấp/Cao (Cha thấp, Mẹ cao)', '−0,45 (−4,47; 3,57)', '4,87 (−1,17; 10,91)', 'n.s.']],
    widths=[4.5, 3.5, 3.5, 2.0])
add_p(doc, 'Ghi chú: *Chỉ tổ hợp "Cha cao / Mẹ thấp" có ý nghĩa trong mô hình đơn biến — gợi ý cha có trình độ cao kỳ vọng con nhiều hơn, đặc biệt khi mẹ không có trình độ tương đương để cân bằng. Tuy nhiên, mối liên quan này biến mất trong mô hình đa biến.', size=9, italic=True)

# ========== 9. THẢO LUẬN ==========
add_page_ref(doc, '284–289', 'J Rural Med', 'Vol. 19(4), 2024')
add_heading(doc, '4. THẢO LUẬN', 2)

add_p(doc, 'Nghiên cứu này khám phá mức độ căng thẳng học tập ở học sinh THCS sử dụng bảng hỏi ESSA với theo dõi ba năm cũng như các yếu tố ảnh hưởng. Kết quả phân tích hồi quy đa biến cho thấy giới tính nữ, số anh chị em, trình độ học vấn của cha, điểm học tập thấp, và học thêm có liên quan đáng kể đến căng thẳng học tập tổng thể và tất cả các lĩnh vực, ngoại trừ "tự kỳ vọng". Hơn nữa, các yếu tố lối sống và mối quan hệ cha mẹ–con có liên quan yếu hoặc không có liên quan đến căng thẳng học tập tổng thể.')

add_p(doc, 'Tăng căng thẳng theo thời gian:', bold=True)
add_p(doc, 'Kết quả cho thấy mức căng thẳng học tập của học sinh tăng đáng kể trong ba năm theo dõi, có thể do nhiều yếu tố. Thứ nhất, khảo sát ban đầu được thực hiện khi học sinh ở năm đầu THCS (lớp 6). Khảo sát theo dõi được thực hiện ba năm sau khi học sinh ở năm cuối THCS (lớp 9). Khi học sinh tiến đến lớp cuối THCS, các em đối mặt với áp lực gia tăng kết hợp với khối lượng công việc quá lớn. Hơn nữa, trong bối cảnh giáo dục Việt Nam, căng thẳng liên quan đến chuẩn bị cho kỳ thi chuyển cấp đặc biệt gay gắt. Xem xét rằng kỳ thi này quan trọng trong việc định hình cuộc sống học tập tương lai, học sinh rất lo lắng và cảm thấy bị áp lực. Tình huống này có thể dẫn đến lo âu, sợ không đủ năng lực, và thiếu tự tin, từ đó khuếch đại mối quan ngại và áp lực. Ngoài ra, chuẩn bị cho các kỳ thi quan trọng thúc đẩy cha mẹ cho con học thêm, tước đi thời gian nghỉ ngơi. Phát hiện của chúng tôi cho thấy tham gia học thêm có tương quan với mức căng thẳng học tập tăng.')

add_p(doc, 'Một yếu tố cần chú ý khác là thời điểm khảo sát theo dõi — năm 2021 trong đại dịch COVID-19 khi các hạn chế cuộc sống hàng ngày đáng kể tại Việt Nam. Một số nghiên cứu đã nêu bật thách thức tâm lý mà học sinh phải đối mặt trong hoạt động học tập trong đại dịch (Ho et al., 2022; Duong et al., 2023), với các biện pháp cách ly, phong tỏa, và chuyển đổi sang học trực tuyến đặt ra thách thức mới có thể tăng mức căng thẳng (Wider et al., 2023).')

add_p(doc, 'Giới tính:', bold=True)
add_p(doc, 'Học sinh nam thể hiện mức căng thẳng học tập thấp hơn nữ, phù hợp với quan sát tương tự trong nghiên cứu tại Hy Lạp ở thanh thiếu niên (Moustaka et al., 2023) và Mỹ (Graves et al., 2021). Sự khác biệt giới trong nhận thức căng thẳng có thể liên quan đến quan sát rằng phụ nữ thường đặt tầm quan trọng lớn hơn vào thành tích học tập. Nữ VTN có xu hướng quan ngại hơn về nhiều khía cạnh cuộc sống và báo cáo trải qua nhiều sự kiện căng thẳng hơn nam (Sun et al., 2013).')

add_p(doc, 'Số anh chị em:', bold=True)
add_p(doc, 'Học sinh có nhiều anh chị em hơn cho thấy căng thẳng học tập tổng thể lớn hơn và căng thẳng cao hơn trong các lĩnh vực "áp lực từ việc học" và "khối lượng công việc". Kết quả này có thể được giải thích bởi xu hướng so sánh thành tích học tập với anh chị em. Áp lực học tập ở học sinh phần nào liên quan đến kỳ vọng của cha mẹ và trường học. Nhận thức rằng cha mẹ có thể ưu ái con có thành tích tốt hơn có thể tạo ra bầu không khí cạnh tranh giữa anh chị em để được cha mẹ công nhận (Danielsbacka & Tanskanen, 2015; Lashewicz & Keating, 2009). Tuy nhiên, một nghiên cứu trước đây phát hiện rằng số anh chị em dự báo ngược chiều căng thẳng giáo dục (Rentala et al., 2019). Ngược lại, một nghiên cứu khác báo cáo rằng thanh thiếu niên có nhiều anh chị em hơn thể hiện mức tự tin cao hơn, có thể giảm căng thẳng (Goel & Aggarwal, 2012).')

add_p(doc, 'Trình độ cha:', bold=True)
add_p(doc, 'Học sinh có cha trình độ cao hơn cho thấy mức căng thẳng cao hơn và căng thẳng cao hơn trong "chán nản" và "khối lượng công việc". Một nghiên cứu tại Ấn Độ chỉ ra rằng nữ VTN có cha mù chữ thể hiện mức căng thẳng cao hơn so với cha biết chữ (Rentala et al., 2019). Các phát hiện tại Mỹ cho thấy VTN từ gia đình có trình độ cha mẹ thấp hơn thể hiện căng thẳng cao hơn — kết quả ngược chiều (Finkelstein et al., 2007). Trong nghiên cứu, không có mối liên quan đáng kể giữa trình độ mẹ và căng thẳng ở con. Một giả thuyết là bản chất nuôi dưỡng và nhân ái thường gắn liền với chăm sóc của mẹ — mẹ thường ưu tiên sức khỏe con và có thể không áp đặt kỳ vọng quá mức (Way et al., 2013). Một nghiên cứu tại TQ báo cáo cha truyền thống đóng vai trò trụ cột và đặt nặng uy tín gia đình, dẫn đến kỳ vọng và áp lực cao hơn lên con (Ji et al., 1993). Sự khác biệt giữa cha và mẹ giải thích tại sao trình độ cha tác động đến căng thẳng con trong khi trình độ mẹ thì không.')

add_p(doc, 'Một mối liên quan thô được quan sát giữa tổ hợp cha trình độ cao + mẹ trình độ thấp và căng thẳng lớn hơn ở con — gợi ý hiện tượng thú vị phản ánh vai trò cha và mẹ trong gia đình bối cảnh Việt Nam. Tuy nhiên, mối liên quan này yếu và biến mất khi điều chỉnh các biến khác.')

add_p(doc, 'Điểm học tập và học thêm:', bold=True)
add_p(doc, 'Điểm học tập và tham gia học thêm có liên quan đến điểm căng thẳng tổng thể. Phát hiện phù hợp với tài liệu: học sinh có điểm trung bình thấp trải qua mức căng thẳng tổng thể cao hơn (Moustaka et al., 2023). Về học thêm, nghiên cứu cho thấy học sinh tham gia lớp học thêm thể hiện mức căng thẳng cao hơn. Phát hiện phù hợp với tài liệu (Moustaka et al., 2023; Sun et al., 2013). Học sinh tham gia lớp học riêng, cá nhân, hoặc ngoại khóa sau giờ học, cuối tuần, và ngày lễ báo cáo trải qua căng thẳng và áp lực. Tương tự, nghiên cứu tại Ấn Độ chỉ ra rằng học thêm đóng góp đáng kể vào căng thẳng ở cả học sinh và giáo viên (Santhi, 2011). Yêu cầu tham dự sáng sớm và tối muộn dẫn đến nhiều vấn đề.')

add_p(doc, 'Điểm thấp hơn và tham gia học thêm liên quan đến mức căng thẳng lớn hơn từ "áp lực từ việc học" và "chán nản" — phù hợp với tài liệu. Học sinh đạt điểm cao hơn trải qua ít lo âu hơn về thành tích, dẫn đến tự tin hơn về giáo dục tương lai. Ngoài ra, học sinh điểm cao (thấp) có xu hướng tự tin cao (thấp) hơn (Tripathy & Srivastava, 2013; Al-Hebaish, 2012).')

add_p(doc, 'Lối sống và thái độ cha mẹ:', bold=True)
add_p(doc, 'Kết quả cho thấy không có mối liên quan giữa các yếu tố lối sống (giờ ngủ, hoạt động thể chất) và căng thẳng tổng thể cũng như hầu hết các lĩnh vực, ngoại trừ mối quan hệ giữa giờ ngủ và "chán nản". Mặc dù không có mối liên quan giữa hoạt động thể chất và căng thẳng, nhiều nghiên cứu cho thấy học sinh tham gia hoạt động thể chất trải qua ít căng thẳng hơn (Moustaka et al., 2023; Pascoe et al., 2020). Tác động giờ ngủ lên "chán nản" được giải thích bởi học sinh ngủ đủ có xu hướng tỉnh táo hơn, dẫn đến tập trung tốt hơn (Boeke et al., 2014; van der Heijden et al., 2018).')

add_p(doc, 'Sự tập trung của cha mẹ cho thấy liên quan đáng kể với "lo lắng về điểm số", trong khi sự chấp nhận của cha mẹ không có liên quan đáng kể. Mặc dù không tìm thấy liên quan giữa tập trung cha mẹ và căng thẳng tổng thể trong mô hình đa biến, tập trung cha mẹ có liên quan đáng kể trong mô hình đơn biến — cho thấy căng thẳng học tập của học sinh phần nào bị ảnh hưởng bởi kỳ vọng và thái độ bảo vệ quá mức của cha mẹ. Một nghiên cứu trước cho thấy nguyên nhân chính của căng thẳng, theo 66% học sinh, là áp lực từ cha mẹ cải thiện thành tích (Deb et al., 2015). Nghiên cứu khác tại Hong Kong cho thấy kỳ vọng cha mẹ đóng góp vào tăng trầm cảm ở học sinh (Ma et al., 2018). Cha mẹ châu Á có thể gây áp lực này vì lo lắng cho sức khỏe con và nhận thức về tính cạnh tranh trong việc nhập học vào các tổ chức uy tín. Kỳ vọng và bảo vệ quá mức được coi là biểu hiện tình yêu và hỗ trợ. Tuy nhiên, sự tập trung của cha mẹ vào ngoại hình hoặc thành tích có thể tác động tiêu cực đến mức căng thẳng (Kang et al., 2017). Do đó, học sinh có cha mẹ áp đặt thường có ít thời gian nghỉ ngơi, vui chơi, và tương tác xã hội — dẫn đến lo âu cao, đặc biệt khi đối mặt với áp lực thi cử.')

add_p(doc, 'Hạn chế nghiên cứu:', bold=True)
add_p(doc, 'Theo hiểu biết tốt nhất của chúng tôi, đây là nghiên cứu dọc ĐẦU TIÊN về căng thẳng học tập và các yếu tố liên quan ở đối tượng Việt Nam tập trung vào học sinh THCS. Thiết kế nghiên cứu cho phép phân tích sự thay đổi điểm số theo thời gian và đưa ra giải thích về mối quan hệ nhân quả. Tuy nhiên, nghiên cứu có một số hạn chế:')
add_p(doc, '(1) Thời điểm khảo sát theo dõi trùng đại dịch COVID-19, gây khó khăn hoàn thành khảo sát theo dõi người tham gia ban đầu.')
add_p(doc, '(2) Thông tin được thu thập bằng đo lường tự báo cáo — có thể thiên lệch do trả lời theo hướng mong muốn xã hội.')
add_p(doc, '(3) Nghiên cứu loại trừ thanh thiếu niên ngoài trường và học sinh trường tư — kết quả chỉ khái quát cho trường công tại VN.')
add_p(doc, '(4) Các yếu tố tiềm ẩn khác có thể ảnh hưởng đến căng thẳng không được đánh giá. Nghiên cứu tương lai nên xem xét tác động của cha mẹ đơn thân, tình trạng kinh tế hộ gia đình, và giao tiếp với bạn bè. Cần nghiên cứu với cỡ mẫu lớn hơn.')

# ========== 10. KẾT LUẬN ==========
add_heading(doc, '5. KẾT LUẬN', 2)
add_p(doc, 'Nghiên cứu thuần tập này đóng góp vào hiểu biết về tình trạng và mối liên quan giữa căng thẳng học tập với các yếu tố khác nhau ở học sinh. Kết quả cho thấy căng thẳng học tập ở học sinh TĂNG trong giai đoạn ba năm nghiên cứu, hàm ý rằng học sinh phải đối mặt với các vấn đề liên quan đến căng thẳng học tập ở năm cuối THCS. Hơn nữa, phát hiện nhấn mạnh tác động của yếu tố gia đình và khối lượng công việc lên căng thẳng học tập. Các chương trình đào tạo và hỗ trợ tại trường học và thúc giục cha mẹ chú ý đến con cái là thiết yếu. Cả giáo viên và cha mẹ có thể hưởng lợi từ đào tạo để nâng cao hiểu biết về con em và phát triển quản lý thời gian hiệu quả cho giáo dục. Điều này sẽ giúp giảm thiểu hoặc ngăn ngừa căng thẳng học tập ở học sinh.')

# ========== 11. TÀI LIỆU THAM KHẢO ==========
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
refs = [
    '1. Ribeiro ÍJS, et al. (2018). Stress and quality of life among university students. Health Prof Educ, 4: 70–77.',
    '2. Deng Y, et al. (2022). Family and academic stress and their impact on students\' depression level. Front Psychiatry, 13: 869337.',
    '3. Pascoe MC, et al. (2020). The impact of stress on students in secondary school and higher education. Int J Adolesc Youth, 25: 104–112.',
    '4. Jiménez-Mijangos LP, et al. (2023). Advances and challenges in the detection of academic stress. Educ Inf Technol, 28: 3637–3666.',
    '5. Klinger DA, et al. (2015). Cross-national trends in perceived school pressure. Eur J Public Health, 25(S2): 51–56.',
    '29. Khanh NĐ, Thi LM, Vinh NA. (2023). Depression in secondary students at Hanoi. Tạp chí Y học VN, 526(1A).',
    '30. Nguyen DT, et al. (2013). Depression, anxiety, and suicidal ideation among Vietnamese secondary school students. BMC Public Health, 13: 1195.',
    '34. Sun J, et al. (2011). ESSA development, validity, and reliability with Chinese students. J Psychoeduc Assess, 29: 534–546.',
    '35. Truc TT, et al. (2015). Validation of ESSA in Vietnam. Asia Pac J Public Health, 27: NP2112–NP2121.',
    '39. Nguyen HTL, et al. (2020). Cyberbullying, parental attitudes, self-harm and suicidal behavior among adolescents in Vietnam. BMC Public Health, 20: 476.',
]
for ref in refs:
    add_p(doc, ref, size=10)
add_p(doc, '(Xem danh mục đầy đủ trong bài gốc — 58 tài liệu tham khảo)', size=10, italic=True)

# ========== 12. BẢNG VIẾT TẮT ==========
add_abbreviation_table(doc, [
    ('ESSA', 'Educational Stress Scale for Adolescents — Thang Căng thẳng Giáo dục dành cho Thanh thiếu niên'),
    ('GSHS', 'Global School-based Student Health Survey — Khảo sát Sức khỏe Học đường Toàn cầu'),
    ('THCS', 'Trung học Cơ sở (Secondary school / Junior high school)'),
    ('THPT', 'Trung học Phổ thông (High school / Senior high school)'),
    ('TMDU', 'Tokyo Medical and Dental University — ĐH Y Nha khoa Tokyo'),
    ('GPA', 'Grade Point Average — Điểm trung bình'),
    ('IRB', 'Institutional Review Board — Hội đồng Đánh giá Thể chế'),
    ('OECD', 'Organisation for Economic Co-operation and Development'),
    ('JSPS', 'Japan Society for the Promotion of Science — Quỹ Khoa học Xã hội Nhật Bản'),
    ('β', 'Beta — Hệ số hồi quy'),
    ('KTC', 'Khoảng tin cậy (Confidence Interval)'),
    ('SD', 'Standard Deviation — Độ lệch chuẩn'),
    ('VTN', 'Vị thành niên (thanh thiếu niên)'),
    ('SKTT', 'Sức khỏe tâm thần'),
    ('ACE', 'Anh chị em (siblings — dùng trong phân tích)'),
])

# ========== 13. PHẢN BIỆN ==========
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')

add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'NC DỌC (longitudinal) 3 năm — ĐẦU TIÊN tại VN về căng thẳng học tập ở HS THCS. Cho phép đánh giá thay đổi theo thời gian và suy luận nhân quả hạn chế. Phù hợp Li 2025 (QT22, dọc 12 tháng Úc) nhưng VƯỢT TRỘI về thời gian theo dõi (3 năm vs 1 năm).',
    'Hue Healthy Adolescent Cohort — hợp tác quốc tế TMDU Nhật Bản + ĐH Y Dược Huế. Thiết kế chặt chẽ: lấy mẫu cụm phân tầng, phỏng vấn mặt đối mặt, phê duyệt đạo đức 2 nước.',
    'ESSA đã xác thực tại VN (Truc et al. 2015, α = 0,83). Cronbach α trong NC này = 0,88 — rất tốt. Phù hợp Vĩnh Lộc 2024 (VN20) cũng dùng ESSA.',
    'HS CẤP 2 (THCS lớp 6→9) — nhóm tuổi ÍT NC nhất tại VN (đa số NC tập trung cấp 3). Đặc biệt phù hợp đề tài (lo âu HS THCS & THPT).',
    'Phát hiện học thêm β = 4,73 — phản ánh bối cảnh VN (thi vào lớp 10 cực kỳ cạnh tranh). Phù hợp Wen 2020 (QT08, áp lực OR = 11,58), Norway 2025 (QT21, bất mãn trường giải thích chính).',
    'Phát hiện cha học vấn cao → con stress cao — phản ánh văn hóa châu Á: cha đóng vai trò "trụ cột", kỳ vọng cao (Ji et al. 1993; Ma et al. 2018). Phù hợp 59 Countries 2025 (QT31, cha mẹ kiểm tra bài AOR = 0,75 bảo vệ nhưng kỳ vọng quá mức tăng nguy cơ).',
    'Open Access (CC BY-NC-ND) — tải miễn phí, dễ truy cập.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'n = 341 phân tích cuối (từ 611 ban đầu — mất 44%) — có thể thiên lệch chọn. Lý do mất mẫu: COVID-19 + thiếu dữ liệu cha mẹ. So: Hoa 2024 (VN01) n = 3.910, Bảo Quyên 2025 (VN16) n = 501 — cắt ngang lớn hơn nhưng KHÔNG DỌC.',
    'ESSA đo CĂNG THẲNG HỌC TẬP — không đo LO ÂU trực tiếp (GAD-7, DASS-21). Liên quan nhưng khác biệt: căng thẳng HT là NGUYÊN NHÂN, lo âu là HẬU QUẢ. Cần kết hợp ESSA + GAD-7 trên cùng mẫu. So: Vĩnh Lộc 2024 (VN20) dùng cả ESSA + DASS-Y.',
    'Follow-up 2021 trùng COVID-19 — khó phân biệt tăng stress do (a) chuyển cấp tự nhiên vs (b) COVID. Hoàng Trung Học 2025 (VN14): lo âu 41,5% → 25,4% sau COVID — gợi ý COVID có tác động nhưng phức tạp.',
    'Chỉ TP Huế (đô thị) — không đại diện nông thôn/DTTS. So: Ngô Anh Vinh 2024 (VN15) DTTS Lạng Sơn 54,4% lo âu; An Giang 2025 (VN18) nông thôn 61,2%.',
    'Tự báo cáo — phỏng vấn mặt đối mặt có thể gây thiên lệch mong muốn xã hội (HS nói giảm căng thẳng trước người phỏng vấn). So: Li 2025 (QT22) dùng online — cũng thiên lệch nhưng theo hướng khác.',
    'KHÔNG đo screen time/MXH — yếu tố quan trọng. Norway 2025 (QT21): MXH giải thích một phần xu hướng. Nature 2025 (QT27): VTN SKTT dùng MXH nhiều hơn (g = 0,46). Zheng 2025: MXH → lo âu β = 0,153.',
    'KHÔNG đo giấc ngủ CHI TIẾT — chỉ 2 nhóm (<8h, ≥8h). Zhu 2025 (QT05): giấc ngủ <5h AOR = 13,71 — yếu tố mạnh nhất nhưng bài này không phát hiện vì phân nhóm thô.',
    'Loại trừ HS ngoài trường + trường tư — nhóm có thể có stress KHÁC (không cạnh tranh thi vào lớp 10). Kết quả chỉ khái quát cho trường THCS công lập đô thị.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống nghiên cứu / Research Gap:', bold=True)
for s in [
    'Cần kết hợp ESSA + GAD-7/DASS-21 trên cùng mẫu — đo CẢ căng thẳng HT và lo âu. So sánh: HS stress cao có nhất thiết lo âu cao? Hoặc stress vừa nhưng lo âu đã nặng? Thảo Vi 2025 (VN19): lạc quan trung gian lo âu → trầm cảm — gợi ý cơ chế phức tạp hơn.',
    'Theo dõi tiếp HS sau thi vào lớp 10 — stress TĂNG HAY GIẢM? Nếu thi xong → giảm (stress thi cử tạm thời), nếu vào trường không ưng → tăng (stress thất vọng). Chưa có NC dọc THCS→THPT tại VN.',
    'Mở rộng đa vùng: Hà Nội (đô thị lớn), TPHCM (đô thị nam), nông thôn (Thanh Hóa — Danh Lâm 2022 VN17), DTTS (Lạng Sơn — Ngô Anh Vinh VN15). So sánh vùng là GAP quan trọng.',
    'Thêm biến screen time + MXH + giấc ngủ chi tiết + bắt nạt vào mô hình — 4 yếu tố mạnh nhất theo y văn quốc tế (Norway QT21, Zhu QT05, Islam QT31, Fassi QT27) nhưng bài này KHÔNG đo.',
    'NC can thiệp dựa trên phát hiện: nếu HỌC THÊM tăng stress (β = 4,73), can thiệp GIẢM HỌC THÊM + tăng thời gian nghỉ có giảm stress không? BMC NMA 2025 (QT29): PE (hoạt động thể chất) SUCRA = 0,51 — có thể thay học thêm bằng thể chất.',
    'Vai trò CHA vs MẸ chi tiết hơn: tổ hợp cha cao/mẹ thấp → stress nhưng biến mất khi điều chỉnh. Cần NC sâu hơn về cơ chế kỳ vọng cha mẹ riêng ở VN — khác biệt văn hóa với phương Tây.',
]:
    add_red(doc, f'• {s}')

# ========== SAVE ==========
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '36_TranThaoVi_2024_JRuralMed.docx')
doc.save(outpath)

# Verify
import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
rels = d.part.rels
img_count = sum(1 for r in rels.values() if 'image' in r.reltype)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}, Images: {img_count}')

# Cross-verify key numbers
import re
checks = ['46,4', '53,5', '341', '611', '2,24', '3,20', '4,73', '−2,85', '−1,79', '0,88', '0,83']
ok = 0
for num in checks:
    if num in t:
        ok += 1
    else:
        print(f'  ! THIEU: {num}')
print(f'  Numbers verified: {ok}/{len(checks)}')
