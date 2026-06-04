# -*- coding: utf-8 -*-
"""Sinh doc tra loi thay: cach dich Journey of the Brave / Happy House /
Power Up-CBTD va cau so lieu ASI-LSAS. 18/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT_DIR = os.path.join(ROOT, '01_Bao-cao')
OUT = os.path.join(OUT_DIR, 'Tra_loi_cach_dich_JourneyOfTheBrave_HappyHouse_PowerUpCBTD_18052026.docx')
PAGE_W = 16.0
RED = RGBColor(0xCC, 0, 0)
BLUE = RGBColor(0x05, 0x63, 0xC1)
ORANGE = RGBColor(0xC0, 0x50, 0x00)

# ---------- helpers ----------
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)

def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(doc, text, bold=False, italic=False, size=12, align='justify', color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def runs_para(doc, parts, size=12, align='justify'):
    """parts: list of (text, dict) — dict keys: bold, italic, color."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.CENTER
    for text, fmt in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size)
        r.bold = fmt.get('bold', False); r.italic = fmt.get('italic', False)
        if 'color' in fmt: r.font.color.rgb = fmt['color']
    return p

def bullet(doc, text, color=None, bold_prefix=None, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = True
        if color is not None: r.font.color.rgb = color
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

def add_hyperlink(paragraph, url, text, size=11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def ref_item(doc, text, doi_label=None, doi_url=None, trans=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    if doi_label and doi_url:
        sep = p.add_run('  ')
        sep.font.name = 'Times New Roman'; sep.font.size = Pt(11)
        add_hyperlink(p, doi_url, doi_label, size=11)
    if trans:
        p2 = doc.add_paragraph()
        p2.paragraph_format.left_indent = Cm(1.0)
        r2 = p2.add_run(trans)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True

def table(doc, headers, rows, widths, hdr_size=10, body_size=10):
    assert sum(widths) <= PAGE_W + 0.05, sum(widths)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(hdr_size)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(body_size)
    return t

# ============================================================
doc = make_doc()

# ---- Tieu de ----
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('TRẢ LỜI: CÁCH DỊCH CÁC TÊN CHƯƠNG TRÌNH CAN THIỆP\n'
                '"Journey of the Brave" — "Happy House" — "Power Up-CBTD"\n'
                'VÀ CÂU SỐ LIỆU ASI / LSAS')
r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tài liệu trả lời câu hỏi — Ngày 18/05/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

P(doc, 'Câu hỏi: Các cụm "Journey of the Brave", "chương trình Happy House", '
       '"Power Up - CBTD" và câu "(Anxiety Sensitivity Index Cohen d = 0,93 tại tuần 6; '
       'Liebowitz Social Anxiety Scale Cohen d = 1,07 tại tuần 12)" nên dịch như thế nào?',
  italic=True, size=11)

# ---- 1. Nguyen tac chung ----
H(doc, '1. Nguyên tắc chung khi dịch tên chương trình can thiệp', 1)
P(doc, 'Ba cụm đầu tiên đều là TÊN RIÊNG của các chương trình can thiệp tâm lý — tức là '
       '"thương hiệu" của chương trình, do nhóm nghiên cứu đặt ra. Tên riêng chương trình '
       'không nên dịch cứng sang tiếng Việt rồi dùng độc lập, vì người đọc sẽ không tra ngược '
       'được về tài liệu gốc. Nguyên tắc xử lý thống nhất như sau:')
bullet(doc, 'Lần đầu xuất hiện: giữ nguyên tên tiếng Anh (đặt trong ngoặc kép), kèm phần '
            'dịch nghĩa tiếng Việt trong ngoặc đơn ngay sau đó.', bold_prefix='Giữ — chú thích: ')
bullet(doc, 'Các lần sau: chỉ cần dùng tên gốc tiếng Anh hoặc tên Việt đã giới thiệu, '
            'không lặp lại cả cụm.', bold_prefix='Lần sau: ')
bullet(doc, 'Phân biệt rõ hai lớp ngôn ngữ: (a) tên thương hiệu chương trình thì GIỮ; '
            '(b) phần mô tả loại hình can thiệp đi kèm (ví dụ "school-based", '
            '"cognitive behavioral group program") thì DỊCH sang tiếng Việt.',
       bold_prefix='Tách lớp: ')
P(doc, 'Cách làm này vừa bảo đảm văn bản đọc thuần Việt, vừa giúp thầy và người đọc '
       'kiểm chứng lại nguồn gốc khi cần.')

# ---- 2. Journey of the Brave ----
H(doc, '2. "Journey of the Brave"', 1)
H(doc, '2.1. Bối cảnh', 2)
P(doc, '"Journey of the Brave" là một chương trình phòng ngừa lo âu phổ quát dựa vào '
       'trường học (universal school-based anxiety prevention programme), xây dựng trên nền '
       'liệu pháp nhận thức – hành vi (Cognitive Behavioural Therapy — CBT). Chương trình do '
       'nhóm của tác giả Yuko Urao (Đại học Chiba, Nhật Bản) phát triển và đưa vào trường học '
       'Nhật từ năm 2014, dành cho trẻ 9–12 tuổi (học sinh các lớp cuối bậc tiểu học). Nội dung '
       'gồm: giáo dục tâm lý, thư giãn, nhận diện và phân tích nỗi lo, tiếp xúc dần với tình '
       'huống gây sợ (exposure), tái cấu trúc nhận thức và rèn kỹ năng quả quyết. Trong kho tài '
       'liệu của thầy, chương trình này tương ứng với hai bài QT068 (Urao và cộng sự, 2018 — '
       'bản 10 buổi) và QT069 (Urao và cộng sự, 2022 — bản 14 buổi ngắn).')

H(doc, '2.2. Phân tích nghĩa', 2)
P(doc, 'Cụm gồm hai từ: "journey" — hành trình, chuyến đi; "the brave" — người dũng cảm. '
       'Tên chương trình là một ẩn dụ: đứa trẻ được hình dung như một người dũng cảm bước vào '
       'một hành trình, trong đó các bài học là từng chặng giúp em đối diện và vượt qua nỗi sợ, '
       'nỗi lo. Đây là cách đặt tên có tính khích lệ, hướng tới trẻ em.')

H(doc, '2.3. Đề xuất dịch', 2)
runs_para(doc, [
    ('Lần đầu: ', {'bold': True}),
    ('chương trình "Journey of the Brave" (Hành trình của người dũng cảm)', {'italic': True}),
    ('.', {}),
])
runs_para(doc, [
    ('Các lần sau: ', {'bold': True}),
    ('"Journey of the Brave"', {'italic': True}),
    (' hoặc ', {}),
    ('chương trình "Hành trình của người dũng cảm".', {'italic': True}),
])
P(doc, 'Các biến thể dịch nghĩa cũng chấp nhận được: "Hành trình của những người dũng cảm" '
       '(số nhiều), hoặc gọn hơn là "Hành trình dũng cảm". Nên tránh các bản dịch mang sắc thái '
       'khẩu ngữ như "Chuyến đi của kẻ gan dạ". Khuyến nghị giữ tên tiếng Anh làm tên chính, vì '
       'cả hai bài QT068 và QT069 cùng dùng tên này — giữ nguyên giúp thầy đối chiếu liền mạch '
       'giữa các tài liệu.')

# ---- 3. Happy House ----
H(doc, '3. "Chương trình Happy House"', 1)
H(doc, '3.1. Bối cảnh', 2)
P(doc, '"Happy House" là chương trình thúc đẩy sức khoẻ tâm thần phổ quát dựa vào trường học '
       'dành cho vị thành niên Việt Nam. Đây chính là phiên bản thích ứng văn hoá Việt Nam của '
       'chương trình RAP (Resourceful Adolescent Program — Chương trình Vị thành niên Giàu '
       'nguồn lực) của Úc. Chương trình được nhóm tác giả Thach Duc Tran và cộng sự (Đại học '
       'Monash, Úc và Đại học Y tế Công cộng Hà Nội) đánh giá qua một thử nghiệm có đối chứng '
       'hai nhánh song song trên 1.084 học sinh lớp 10 tại 8 trường trung học phổ thông ở Hà Nội, '
       'công bố năm 2023. Trong kho tài liệu của thầy, bài này tương ứng mã VN030.')

H(doc, '3.2. Phân tích nghĩa', 2)
P(doc, 'Cụm gồm: "happy" — hạnh phúc, vui vẻ; "house" — ngôi nhà. Tên gợi hình ảnh một "ngôi '
       'nhà" an toàn, ấm áp về mặt cảm xúc — nơi vị thành niên được trang bị kỹ năng để giữ '
       'tinh thần khoẻ mạnh.')

H(doc, '3.3. Đề xuất dịch', 2)
runs_para(doc, [
    ('Lần đầu: ', {'bold': True}),
    ('chương trình "Happy House" (Ngôi nhà Hạnh phúc)', {'italic': True}),
    ('.', {}),
])
runs_para(doc, [
    ('Các lần sau: ', {'bold': True}),
    ('"Happy House"', {'italic': True}),
    (' hoặc ', {}),
    ('chương trình "Ngôi nhà Hạnh phúc".', {'italic': True}),
])
P(doc, 'Lưu ý quan trọng: đây là chương trình được thực hiện ngay TẠI VIỆT NAM, do nhóm '
       'nghiên cứu Việt – Úc đặt tên. "Happy House" là tên thương hiệu mà chính nhóm này đặt '
       'cho phiên bản RAP đã Việt hoá. Nếu trong tài liệu gốc hoặc tài liệu truyền thông của dự '
       'án, nhóm có sử dụng một tên tiếng Việt chính thức, thì NÊN ưu tiên dùng đúng tên tiếng '
       'Việt của nhóm. Nếu không có, "Ngôi nhà Hạnh phúc" là bản dịch tự nhiên và sát nghĩa '
       'nhất. Nên dịch "house" là "ngôi nhà" (đầy đủ, ấm áp) chứ không dịch trống thành "nhà".')

# ---- 4. Power Up-CBTD ----
H(doc, '4. "Power Up – CBTD"', 1)
H(doc, '4.1. Bối cảnh', 2)
P(doc, '"Power Up-CBTD" là một chương trình phòng ngừa trầm cảm ngắn, dựa vào trường học, '
       'theo hình thức trị liệu nhóm nhận thức – hành vi (Cognitive Behavioral Group Therapy '
       '— CBGT). Chương trình do nhóm tác giả Qianyun He và cộng sự (Đại học Sư phạm Bắc Kinh, '
       'Trung Quốc) phát triển và đánh giá qua một thử nghiệm có đối chứng ngẫu nhiên quy mô '
       'thử nghiệm (pilot RCT) trên 87 học sinh các lớp 5, 6 và 9 tại tỉnh Hà Nam, Trung Quốc. '
       'Trong kho tài liệu của thầy, bài này tương ứng mã QT073.')

H(doc, '4.2. Xác minh tên đầy đủ của "CBTD"', 2)
runs_para(doc, [
    ('Đã đối chiếu trực tiếp với bản PDF gốc (trang 1, phần Tóm tắt). Tên đầy đủ chính thức '
     'của chương trình là: ', {}),
    ('"Power Up: Cognitive Behavioral Therapy for Depressive Symptoms (Power Up-CBTD)"', {'italic': True}),
    ('.', {}),
])
P(doc, 'Như vậy, "CBTD" KHÔNG đơn thuần là "CBT for Depression". Phần "CBTD" là viết tắt của '
       '"Cognitive Behavioral Therapy for Depressive Symptoms" — nghĩa là "Liệu pháp nhận thức '
       '– hành vi cho các triệu chứng trầm cảm". Đây là một điểm cần đặc biệt cẩn trọng: nếu '
       'suy đoán acronym sẽ dễ dịch sai thành "trầm cảm" thay vì "triệu chứng trầm cảm".')

H(doc, '4.3. Phân tích nghĩa "Power Up"', 2)
P(doc, '"Power up" là một ẩn dụ quen thuộc từ trò chơi điện tử — chỉ vật phẩm hoặc khoảnh '
       'khắc giúp nhân vật "được tiếp thêm sức mạnh", mạnh lên. Khi dùng làm tên chương trình, '
       'nó gợi ý việc "tiếp sức", "nạp năng lượng tinh thần" cho học sinh. Vì là tên thương '
       'hiệu, nên giữ nguyên "Power Up" và không dịch cứng thành "Tăng lực".')

H(doc, '4.4. Đề xuất dịch', 2)
runs_para(doc, [
    ('Lần đầu: ', {'bold': True}),
    ('chương trình "Power Up-CBTD" (Power Up: Liệu pháp nhận thức – hành vi cho các triệu '
     'chứng trầm cảm — Cognitive Behavioral Therapy for Depressive Symptoms)', {'italic': True}),
    ('.', {}),
])
runs_para(doc, [
    ('Các lần sau: ', {'bold': True}),
    ('"Power Up-CBTD".', {'italic': True}),
])
P(doc, 'Nếu cần diễn giải nghĩa "Power Up" trong ngoặc cho người đọc phổ thông, có thể thêm '
       '"tạm hiểu: tiếp thêm sức mạnh tinh thần" — nhưng đây chỉ là chú thích nghĩa, không thay '
       'thế tên thương hiệu.')

# canh bao dinh chinh
P(doc, '⚠ Cảnh báo cần đính chính (xem chi tiết ở mục 7 và phần Tham khảo): trong một file '
       'nháp nội bộ của dự án, trích dẫn bài He 2025 đang bị ghi SAI tên bài, danh sách tác giả '
       'và số tập – số trang. Cần sửa lại theo bản gốc đã xác minh trước khi đưa vào bài viết '
       'chính thức.', color=RED)

# ---- 5. Cau so lieu ASI/LSAS ----
H(doc, '5. Câu số liệu "(Anxiety Sensitivity Index... ; Liebowitz Social Anxiety Scale...)"', 1)
H(doc, '5.1. Nguồn của câu này', 2)
P(doc, 'Câu số liệu "(Anxiety Sensitivity Index Cohen d = 0,93 tại tuần 6; Liebowitz Social '
       'Anxiety Scale Cohen d = 1,07 tại tuần 12)" lấy từ nghiên cứu của Bress và cộng sự '
       '(2024) đăng trên JAMA Network Open — thử nghiệm lâm sàng ngẫu nhiên đánh giá ứng dụng '
       'di động "Maya" cho thanh niên 18–25 tuổi mắc rối loạn lo âu. Đây là nghiên cứu khác '
       'với ba chương trình ở các mục trên; trong kho tài liệu của thầy, bài này tương ứng mã '
       'B3. Số liệu đã được đối chiếu nguyên văn với văn bản gốc của bài báo.')

H(doc, '5.2. Đề xuất câu dịch hoàn chỉnh', 2)
runs_para(doc, [
    ('"(Chỉ số Nhạy cảm Lo âu — Anxiety Sensitivity Index, ASI — đạt hệ số d của Cohen = '
     '0,93 ở tuần thứ 6; Thang Đánh giá Lo âu Xã hội Liebowitz — Liebowitz Social Anxiety '
     'Scale, LSAS — đạt hệ số d của Cohen = 1,07 ở tuần thứ 12)".', {'italic': True}),
])

H(doc, '5.3. Giải thích từng thành phần', 2)
table(doc,
    ['Cụm gốc', 'Đề xuất dịch', 'Giải thích'],
    [
        ['Anxiety Sensitivity Index (ASI)',
         'Chỉ số Nhạy cảm Lo âu (Anxiety Sensitivity Index — ASI)',
         'Thang đo mức độ một người sợ hãi chính các cảm giác cơ thể của lo âu (tim đập '
         'nhanh, khó thở, chóng mặt...) vì cho rằng các cảm giác đó nguy hiểm. Giữ tên tiếng '
         'Anh và viết tắt ASI khi nhắc lần đầu.'],
        ['Liebowitz Social Anxiety Scale (LSAS)',
         'Thang Đánh giá Lo âu Xã hội Liebowitz (Liebowitz Social Anxiety Scale — LSAS)',
         '"Liebowitz" là họ của tác giả thang đo (Michael R. Liebowitz) — giữ nguyên, không '
         'dịch, không phiên âm. Thang đo mức độ sợ hãi và né tránh trong các tình huống xã hội. '
         'Giữ viết tắt LSAS.'],
        ['Cohen d',
         'hệ số d của Cohen (Cohen’s d)',
         'Chỉ số đo kích thước hiệu ứng (effect size). Ký hiệu "d" giữ nguyên, KHÔNG dịch. '
         'Theo chuẩn Cohen: d ≈ 0,2 là hiệu quả nhỏ; ≈ 0,5 trung bình; ≥ 0,8 là lớn.'],
        ['= 0,93 / = 1,07',
         '= 0,93 / = 1,07',
         'Giữ dấu phẩy thập phân theo quy ước tiếng Việt. Cả hai giá trị đều ≥ 0,8 nên '
         'đều thuộc mức hiệu quả LỚN.'],
        ['tại tuần 6 / tại tuần 12',
         'ở tuần thứ 6 (kết thúc can thiệp) / ở tuần thứ 12 (theo dõi sau can thiệp)',
         'Trong nghiên cứu gốc, tuần 6 là thời điểm KẾT THÚC can thiệp (end point); tuần 12 '
         'là thời điểm THEO DÕI sau can thiệp 6 tuần (follow-up). Nên ghi rõ để câu dịch không '
         'gây hiểu nhầm.'],
    ],
    widths=[3.7, 4.6, 7.7], hdr_size=10, body_size=9)

H(doc, '5.4. Lưu ý kỹ thuật khi dịch câu này', 2)
P(doc, 'Hai hệ số d nói trên là kích thước hiệu ứng của thay đổi THEO THỜI GIAN trong toàn '
       'bộ mẫu (so sánh điểm số trước – sau can thiệp), KHÔNG phải mức chênh lệch so với một '
       'nhóm chứng không can thiệp. Nghiên cứu Maya chia người tham gia thành 3 nhánh khích lệ '
       'khác nhau và không tìm thấy khác biệt giữa các nhánh. Vì vậy, khi dịch và diễn giải, '
       'nên dùng cách nói "mức cải thiện theo thời gian" và tránh khẳng định "hiệu quả vượt '
       'trội so với nhóm chứng" — điều này sẽ phản ánh đúng thiết kế nghiên cứu gốc.')

# ---- 6. Bang tong hop ----
H(doc, '6. Bảng tổng hợp đề xuất dịch', 1)
table(doc,
    ['Cụm gốc', 'Bản chất', 'Đề xuất dịch (lần đầu xuất hiện)', 'Dùng các lần sau'],
    [
        ['Journey of the Brave', 'Tên chương trình (Nhật Bản)',
         'chương trình "Journey of the Brave" (Hành trình của người dũng cảm)',
         '"Journey of the Brave" / chương trình "Hành trình của người dũng cảm"'],
        ['Happy House', 'Tên chương trình (Việt Nam — bản Việt hoá của RAP)',
         'chương trình "Happy House" (Ngôi nhà Hạnh phúc)',
         '"Happy House" / "Ngôi nhà Hạnh phúc"'],
        ['Power Up-CBTD', 'Tên chương trình (Trung Quốc)',
         'chương trình "Power Up-CBTD" (Power Up: Liệu pháp nhận thức – hành vi cho các '
         'triệu chứng trầm cảm)',
         '"Power Up-CBTD"'],
        ['Anxiety Sensitivity Index', 'Tên thang đo',
         'Chỉ số Nhạy cảm Lo âu (Anxiety Sensitivity Index — ASI)', 'ASI'],
        ['Liebowitz Social Anxiety Scale', 'Tên thang đo',
         'Thang Đánh giá Lo âu Xã hội Liebowitz (Liebowitz Social Anxiety Scale — LSAS)',
         'LSAS'],
        ['Cohen d', 'Chỉ số thống kê',
         'hệ số d của Cohen (Cohen’s d)', 'hệ số d / d'],
    ],
    widths=[2.9, 2.7, 6.9, 3.5], hdr_size=10, body_size=9)

# ---- 7. Phan bien ----
H(doc, '7. Một số lưu ý và quan điểm phản biện', 1)
P(doc, '(Phần dưới đây in màu đỏ theo quy ước tài liệu — là các lưu ý cần cân nhắc khi sử '
       'dụng các bản dịch trên.)', italic=True, size=11)

bullet(doc, 'Tên thương hiệu chương trình nên được giữ nguyên tiếng Anh kèm chú thích, '
            'không thay thế hoàn toàn bằng bản dịch. Lý do: ba chương trình này đến từ ba quốc '
            'gia (Nhật, Việt Nam, Trung Quốc); nếu chỉ dùng tên Việt, người đọc rất khó tra '
            'ngược về tài liệu gốc để kiểm chứng.', color=RED,
       bold_prefix='Giữ tên gốc. ')
bullet(doc, 'Đối chiếu trực tiếp với bản PDF gốc (QT073, trang 1) cho thấy "CBTD" = '
            '"Cognitive Behavioral Therapy for Depressive Symptoms". Một bản nháp nội bộ của '
            'dự án (file paper-may/_workspace/bai2_dump.txt, dòng 68) đang trích dẫn bài He '
            '2025 SAI: ghi tên bài là "Effectiveness of an 8-session CBT-based prevention '
            'program... J Affect Disord 367, 412–421" với tác giả "He X., Liu Y., Zhou J., '
            'Wang Q." Bản gốc đúng là: He Q., Li J., Wang J., Qu Z. (2026), "Preventing '
            'depression in Chinese children and adolescents...", J Affect Disord 394:120559. '
            'Em chưa tự sửa file nháp này; đề nghị thầy xác nhận để em đính chính.', color=RED,
       bold_prefix='Đính chính trích dẫn. ')
bullet(doc, 'Hai hệ số d = 0,93 và d = 1,07 là kích thước hiệu ứng trước – sau trong toàn '
            'mẫu, không phải hiệu quả so với nhóm chứng. Khi dịch, nên diễn đạt là "mức cải '
            'thiện theo thời gian" để không phóng đại kết quả.', color=RED,
       bold_prefix='Phân biệt loại hiệu ứng. ')
bullet(doc, 'Cụm "tại tuần 6" và "tại tuần 12" nên được dịch kèm chú thích "kết thúc can '
            'thiệp" và "theo dõi sau can thiệp", vì bản thân con số tuần không nói rõ đó là '
            'thời điểm nào trong tiến trình nghiên cứu.', color=RED,
       bold_prefix='Ghi rõ mốc thời gian. ')
bullet(doc, 'Câu số liệu ASI/LSAS thuộc về nghiên cứu ứng dụng Maya (Bress 2024) — một '
            'nghiên cứu KHÁC với ba chương trình Journey of the Brave, Happy House và Power '
            'Up-CBTD. Nếu trong bài viết của thầy bốn nội dung này nằm cạnh nhau, nên tách bạch '
            'rõ để tránh người đọc gán nhầm số liệu cho sai chương trình.', color=RED,
       bold_prefix='Tránh gán nhầm nguồn. ')

# ---- 8. Tham khao ----
H(doc, '8. Tham khảo & truy vết', 1)
P(doc, 'Nguồn bên ngoài (có thể kiểm chứng):', bold=True, size=11)

ref_item(doc,
    'Bress JN, Falk A, Schier MM, và cộng sự (2024). "Efficacy of a Mobile App-Based '
    'Intervention for Young Adults With Anxiety Disorders: A Randomized Clinical Trial." '
    'JAMA Network Open, 7(8):e2428372. — Nguồn gốc của câu số liệu ASI/LSAS (ứng dụng Maya).',
    doi_label='doi:10.1001/jamanetworkopen.2024.28372',
    doi_url='https://doi.org/10.1001/jamanetworkopen.2024.28372',
    trans='Tạm dịch: "Hiệu quả của một can thiệp dựa trên ứng dụng di động cho thanh niên '
          'mắc rối loạn lo âu: Thử nghiệm lâm sàng ngẫu nhiên".')

ref_item(doc,
    'He Q, Li J, Wang J, Qu Z (2026). "Preventing depression in Chinese children and '
    'adolescents: A pilot study of a brief school-based cognitive behavioral group program." '
    'Journal of Affective Disorders, 394:120559. — Nguồn xác minh tên đầy đủ "Power Up: '
    'Cognitive Behavioral Therapy for Depressive Symptoms (Power Up-CBTD)".',
    doi_label='doi:10.1016/j.jad.2025.120559',
    doi_url='https://doi.org/10.1016/j.jad.2025.120559',
    trans='Tạm dịch: "Phòng ngừa trầm cảm ở trẻ em và vị thành niên Trung Quốc: Nghiên cứu '
          'thử nghiệm về một chương trình trị liệu nhóm nhận thức – hành vi ngắn dựa vào '
          'trường học".')

ref_item(doc,
    'Tran TD, Nguyen H, Shochet I, và cộng sự (2023). "School-based universal mental health '
    'promotion intervention for adolescents in Vietnam: Two-arm, parallel, controlled trial." '
    'Cambridge Prisms: Global Mental Health, 10:e69. PMCID: PMC10643236. — Nguồn của chương '
    'trình "Happy House".',
    doi_label='doi:10.1017/gmh.2023.66',
    doi_url='https://doi.org/10.1017/gmh.2023.66')

ref_item(doc,
    'Urao Y và cộng sự (2018). "Effectiveness of a cognitive behavioural therapy-based '
    'anxiety prevention programme at an elementary school in Japan: a quasi-experimental '
    'study." Child and Adolescent Psychiatry and Mental Health, 12:33. PMCID: PMC6007075. '
    '— QT068, chương trình "Journey of the Brave" (bản 10 buổi).',
    doi_label='doi:10.1186/s13034-018-0240-5',
    doi_url='https://doi.org/10.1186/s13034-018-0240-5')

ref_item(doc,
    'Urao Y và cộng sự (2022). "School-based cognitive behavioural intervention programme '
    'for addressing anxiety in 10- to 11-year-olds using short classroom activities in Japan: '
    'a quasi-experimental study." BMC Psychiatry, 22:658. PMCID: PMC9594947. — QT069, '
    'chương trình "Journey of the Brave" (bản 14 buổi ngắn).',
    doi_label='doi:10.1186/s12888-022-04326-y',
    doi_url='https://doi.org/10.1186/s12888-022-04326-y')

ref_item(doc,
    'Liebowitz MR (1987). "Social Phobia." Modern Problems of Pharmacopsychiatry, 22:141-173. '
    'PMID: 2885745. — Nguồn gốc của thang đo LSAS.',
    doi_label='doi:10.1159/000414022',
    doi_url='https://doi.org/10.1159/000414022')

ref_item(doc,
    'Taylor S, Zvolensky MJ, Cox BJ, và cộng sự (2007). "Robust dimensions of anxiety '
    'sensitivity: development and initial validation of the Anxiety Sensitivity Index-3." '
    'Psychological Assessment, 19(2):176-188. — Bản kiểm định thang ASI-3. (Khái niệm ASI gốc: '
    'Reiss S, Peterson RA, Gursky DM, McNally RJ, 1986, Behaviour Research and Therapy, '
    '24(1):1-8.)',
    doi_label='doi:10.1037/1040-3590.19.2.176',
    doi_url='https://doi.org/10.1037/1040-3590.19.2.176')

ref_item(doc,
    'Cohen J (1988). Statistical Power Analysis for the Behavioral Sciences. 2nd ed. '
    'Hillsdale, NJ: Lawrence Erlbaum Associates. — Chuẩn diễn giải hệ số d (0,2 nhỏ; 0,5 '
    'trung bình; 0,8 lớn).')

P(doc, 'Truy vết nội bộ trong kho tài liệu của dự án:', bold=True, size=11)
bullet(doc, '02_Papers-goc/canonical_index.json — các mã QT068, QT069 (descriptor '
            '"Urao_2022_JourneyBrave_14sessions_BMC") và QT073 ("He_2025_PowerUpCBTD_China_'
            'JAD").', size=11)
bullet(doc, '06_Scripts/add_4_papers_to_canonical.py (dòng 50–82) — tóm tắt nội bộ QT068 và '
            'QT069 (Journey of the Brave).', size=11)
bullet(doc, '06_Scripts/add_QT072_QT073_canonical.py (dòng 61–83) — tóm tắt nội bộ QT073 '
            '(Power Up-CBTD).', size=11)
bullet(doc, '06_Scripts/dich_VN030_HappyHouse.py — bản dịch đầy đủ VN030 (Happy House, '
            'Tran và cộng sự 2023).', size=11)
bullet(doc, '03_Ban-dich/B3_JAMA_App_raw.txt (dòng 327–336) — văn bản gốc nghiên cứu Maya; '
            'chứa nguyên văn các hệ số d của ASI và LSAS.', size=11)
bullet(doc, '02_Papers-goc/QT073_He_2025_PowerUpCBTD_China_JAD.pdf (trang 1) — nguồn đối '
            'chiếu trực tiếp để xác minh "CBTD" = "Cognitive Behavioral Therapy for '
            'Depressive Symptoms".', size=11)
bullet(doc, 'Đồ thị tri thức (KG): các nút PAPER_QT068_Urao_2018, PAPER_QT069_Urao_2022 và '
            'khái niệm CONCEPT_JOURNEY_OF_THE_BRAVE — định nghĩa trong '
            '06_Scripts/add_4_papers_to_canonical.py (dòng 151–175).', size=11)

P(doc, 'Đính chính file dự án cần thực hiện:', bold=True, size=11, color=RED)
bullet(doc, 'paper-may/_workspace/bai2_dump.txt — dòng 68: trích dẫn He 2025 ghi sai tên bài, '
            'danh sách tác giả và số tập – số trang. Cần sửa theo bản gốc đã xác minh ở mục '
            'tham khảo số 2. EM CHƯA TỰ SỬA — chờ thầy/anh xác nhận vì đây là file nháp.',
       size=11, color=RED)
bullet(doc, '06_Scripts/dich_VN030_HappyHouse.py — dòng 102: ghi số định danh bài Happy House '
            'là "e58"; đối chiếu bản gốc (PMC10643236, Cambridge Core) cho thấy số đúng là '
            '"e69" (DOI 10.1017/gmh.2023.66 vẫn đúng). EM CHƯA TỰ SỬA — chờ thầy/anh xác nhận.',
       size=11, color=RED)

doc.save(OUT)
print('DA LUU:', OUT)
print('So doan:', len(doc.paragraphs))
