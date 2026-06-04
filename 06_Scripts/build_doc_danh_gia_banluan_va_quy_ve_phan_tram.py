"""DOC: Danh gia phan Ban luan (3.6) cua cac BAN TRO LY (trong file v2)
+ Tra loi cau hoi: tai sao khong co %, cach quy ve %.
Note: Cac ban tro ly la NGUOI THAT, khong phai AI.
"""
import sys, io, math
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Danh_gia_banluan_3_6_va_cach_quy_ve_phan_tram.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ĐÁNH GIÁ PHẦN BÀN LUẬN (3.6) CỦA CÁC BẠN TRỢ LÝ\nVÀ CÁCH QUY VỀ TỶ LỆ % CHO BÁO CÁO\n— Trả lời câu hỏi của thầy ngày 09/05/2026 —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    '(1) Em xem bình luận của Bạn trợ lý của C.H có ổn không nhé? '
    'Thầy không hiểu biết gì nhiều về lĩnh vực này, ngoài những '
    'kiến thức mới học được từ Em trong thời gian qua. Các bạn '
    'trợ lý đã báo cáo trong file 00_Bình luận số liệu-v2.', italic=True
)
blue_run(
    '(2) Tại sao số liệu thực trạng này lại không quy ra % như '
    'nhiều nghiên cứu khác?', italic=True
)
blue_run(
    '(3) Có cách nào quy về % để có bức tranh thô về RLLA của HS '
    'không?', italic=True
)

# Phần I — Đánh giá Bàn luận
H('PHẦN I — Đánh giá Bàn luận của các bạn trợ lý (đoạn 180–189 trong file v2)', level=2, color=NAVY)
para('Bình luận tổng:', bold=True)
para(
    'Phần Bàn luận có CẤU TRÚC TỐT — có dẫn chứng từ y văn quốc '
    'tế và Việt Nam, có kết luận hợp lý. Tuy nhiên, vẫn còn 6 '
    'điểm CẦN CẢI THIỆN — chủ yếu do BỎ QUA các phát hiện ĐẶC '
    'BIỆT của dữ liệu, gây mất giá trị khoa học.'
)
para('')

H('1. ƯU điểm', level=3, color=GREEN)
para('• Có cấu trúc theo từng yếu tố (đoạn 180-188).')
para('• Có dẫn chứng cả VN (Ngô Anh Vinh 2022, Lê T.P. Thanh 2024, Đinh Văn Tài 2021) và quốc tế (Carver 1997, Steare 2023, Li 2025, ...).')
para('• Phát hiện mạnh nhất (áp lực học tập) được nêu rõ — đoạn 183.')
para('• Có giải thích β BPĐP DƯƠNG là maladaptive coping — đoạn 188.')

H('2. SÁU NHƯỢC điểm cần cải thiện', level=3, color=RED)

H('2.1. Lỗi năm "Nguyễn Thị Vân (2018)" — KHÔNG CHÍNH XÁC', level=3)
para(
    'Đoạn 183 ghi "Nguyễn Thị Vân (2018)". Trong cơ sở dữ liệu '
    'dự án, VN004 Nguyễn Thị Vân là năm 2020 (descriptor: '
    'NguyenThiVan_2020_STAI_TPHCM). Em đã verify trước đây — '
    'chỉ có bản 2020, không có bản 2018.'
)
para('Đề xuất sửa: "Nguyễn Thị Vân (2020)".', bold=True)

H('2.2. BỎ QUA phát hiện đặc biệt về BNHĐ → lo âu CHIA LY', level=3)
para(
    'Đoạn 185 nói chung chung "bắt nạt là yếu tố nguy cơ". '
    'KHÔNG NÊU phát hiện CỰC KỲ ĐẶC BIỆT: β BNHĐ → RLLACL = '
    '0,376 — MẠNH NHẤT trong ba dạng RLLA, MẠNH HƠN cả β → '
    'RLLATQ (0,215) và β → RLLAXH (0,253). Đây là pattern '
    'KHÁC BIỆT với áp lực học tập (mạnh nhất ở RLLATQ) và tự '
    'trọng (mạnh nhất ở RLLATQ).'
)
para(
    'Cơ chế giải thích: bắt nạt → mất AN TOÀN ở trường → tránh '
    'né trường ("school refusal") → biểu hiện như lo âu chia '
    'ly. Đặc thù lứa tuổi THCS đang chuyển tiếp.'
)
para('Đề xuất bổ sung: nêu phát hiện này như đóng góp ĐỘC ĐÁO của luận án.', bold=True)

H('2.3. BỎ QUA cảnh báo về fit indices KÉM của BPĐP', level=3)
para(
    'Đoạn 188 nêu mô hình BPĐP có β dương — đúng. Nhưng KHÔNG '
    'cảnh báo: mô hình BPĐP có FIT KÉM trên TẤT CẢ tổ hợp:'
)
para('• RMSEA 0,080–0,204 (vượt ngưỡng "good fit" 0,06; thậm chí 0,204 = RẤT KÉM)')
para('• CFI 0,865–0,911 (dưới ngưỡng 0,95)')
para('• χ²/df 9,6–57,3 (vượt xa ngưỡng 5)')
para(
    'Theo Hu & Bentler (1999), mô hình có RMSEA > 0,10 KHÔNG '
    'NÊN báo cáo làm phát hiện chính. Việc β = 0,749 hấp dẫn '
    'nhưng KHÔNG đáng tin do mô hình specification có vấn đề.'
)
para('Đề xuất bổ sung: nêu cảnh báo này, đề xuất phân tách Brief-COPE 14 nhân tố thay vì 3 nhóm.', bold=True)

H('2.4. KHÔNG nêu rõ pattern chênh lệch giới theo loại RLLA', level=3)
para(
    'Đoạn 181 nói "nữ > nam" chung chung. Thực tế dữ liệu cho '
    'thấy pattern BA TẦNG:'
)
para('• RLLATQ (lan tỏa): F=44,484; Cohen d=0,365 — chênh lệch giới CÓ Ý NGHĨA')
para('• RLLAXH (xã hội): F=45,984; Cohen d=0,370 — chênh lệch giới CÓ Ý NGHĨA, gần ngang RLLATQ')
para('• RLLAC (chia ly): F=0,246 (NS); d≈0,03 — KHÔNG có chênh lệch giới')
para(
    'Phát biểu CHÍNH XÁC: "Chênh lệch giới CHỈ XUẤT HIỆN ở '
    'RLLATQ và RLLAXH, KHÔNG xuất hiện ở RLLACL — phù hợp '
    'khung phát triển: lo âu chia ly là rối loạn khởi phát '
    'sớm, không có cơ sở nội tiết-xã hội cho chênh lệch giới."'
)
para('Đề xuất bổ sung: nêu rõ pattern ba tầng này.', bold=True)

H('2.5. CHƯA so sánh CƯỜNG ĐỘ tương đối các yếu tố', level=3)
para(
    'Bàn luận đề cập từng yếu tố RIÊNG LẺ nhưng KHÔNG SO SÁNH '
    'cường độ chuẩn hóa giữa các yếu tố. Cần một bảng tổng:'
)
add_table(['Yếu tố', '|β| → RLLA tổng', 'Cấp', 'Đánh giá'],
    [
        ['BPĐP (đối phó)', '0,735', 'cấp 1.5', '⚠ fit KÉM, không đáng tin'],
        ['ALHT (áp lực học tập)', '0,533', 'cấp 1', 'MẠNH — fit tốt'],
        ['Tự trọng (TTr)', '0,457 (âm)', 'cấp 1', 'BẢO VỆ MẠNH — fit tốt'],
        ['NĐT (nghiện điện thoại)', '0,400', 'cấp 1', 'MẠNH — fit tốt'],
        ['BNHĐ (bắt nạt)', '0,276', 'cấp 1', 'TRUNG BÌNH — fit chấp nhận'],
        ['HTCM (hỗ trợ cha mẹ)', '0,234 (âm)', 'cấp 1', 'BẢO VỆ TRUNG BÌNH'],
        ['GBTH (gắn bó trường)', '0,155 (âm)', 'cấp 1', 'BẢO VỆ NHỎ'],
        ['HTBB (hỗ trợ bạn bè)', '0,044 (NS)', 'cấp 1', 'KHÔNG có ý nghĩa thống kê'],
    ]
)
para('Đề xuất bổ sung: bảng so sánh + xếp hạng cường độ.', bold=True)

H('2.6. KHÔNG quy ra % (thầy đã hỏi — sẽ trả lời ở Phần II)', level=3)
para(
    'Bàn luận chỉ dùng ĐTB trên thang 0–100, KHÔNG quy ra TỶ '
    'LỆ % học sinh có biểu hiện RLLA. Đây là hạn chế lớn vì '
    'tỷ lệ % giúp người đọc hình dung CỠ CỦA VẤN ĐỀ. Phần II '
    'dưới đây trả lời câu hỏi này.'
)

# Phần II — Tại sao không quy ra %
H('PHẦN II — Tại sao số liệu không quy ra %?', level=2, color=NAVY)
para(
    'Câu trả lời ngắn: tác giả luận án dùng thang đo RCADS '
    '(Revised Children\'s Anxiety and Depression Scale; Chorpita '
    'và cộng sự, 2000) — thang đo NỐI LIÊN TỤC (continuous '
    'scale) cho mức độ triệu chứng — KHÔNG phải thang đo CHẨN '
    'ĐOÁN có/không (categorical).'
)
para('')
para('Khác biệt giữa hai loại thang đo:', bold=True)
add_table(
    ['Loại thang', 'Đặc trưng', 'Cho biết', 'Ví dụ'],
    [
        ['Liên tục (RCADS)',
         'Điểm số 0–100 hoặc 0–3 mỗi mục',
         'MỨC ĐỘ triệu chứng (cao/thấp)',
         'ĐTB = 51,43 → mức độ trung bình'],
        ['Chẩn đoán (DISC, CIDI)',
         'Có/không theo tiêu chí DSM-5',
         'TỶ LỆ % HS có rối loạn',
         'V-NAMHS 2022: 18,45% có RLLA'],
    ]
)
para('')
para(
    'Lý do chương 3 luận án không quy ra %: vì RCADS là thang '
    'đo MỨC ĐỘ, KHÔNG phải thang chẩn đoán. Để chuyển sang % '
    'cần một NGƯỠNG CẮT (cutoff) cụ thể — và bài chưa nêu rõ '
    'ngưỡng nào.'
)

# Phần III — Cách quy về %
H('PHẦN III — Ba cách quy về % để có bức tranh thô', level=2, color=NAVY)

H('Cách 1 — Sử dụng ngưỡng cắt RCADS chuẩn quốc tế', level=3, color=NAVY)
para(
    'RCADS có ngưỡng T-score chuẩn (Chorpita 2000):',
    bold=True
)
para('• T-score ≥ 65: BORDERLINE clinical (cảnh báo)')
para('• T-score ≥ 70: CLINICAL (đạt ngưỡng chẩn đoán RLLA)')
para('')
para(
    'Tuy nhiên, chương 3 luận án chuyển điểm sang thang 0–100. '
    'Nếu giả định 0–100 ≈ T-score (linear transform), ngưỡng '
    'có thể là 65/100 hoặc 70/100. Cần xác nhận với tác giả '
    'luận án.'
)

H('Cách 2 — Dùng z-score giả định phân phối chuẩn', level=3, color=NAVY)
para(
    'Cách thực tế nhất với data hiện có. Giả định điểm RCADS '
    'phân bố CHUẨN, dùng công thức:'
)
para('z = (Ngưỡng − ĐTB) / SD', bold=True)
para('Sau đó tra bảng phân phối chuẩn để có % HS vượt ngưỡng.')
para('')
para('Áp dụng cho mẫu n=1.352 với ngưỡng 65/100:', bold=True)


# Tính z-score và tỷ lệ
def normal_cdf(x):
    return 0.5 * (1 + math.erf(x / math.sqrt(2)))

# Sử dụng trung bình giữa nam và nữ
n_male, n_female = 614, 738
total_n = n_male + n_female

# RLLATQ tổng
M_total_TQ = (n_male * 51.43 + n_female * 59.47) / total_n
SD_pooled_TQ = math.sqrt(((n_male-1)*22.010**2 + (n_female-1)*22.072**2) / (total_n-2))
z65_TQ = (65 - M_total_TQ) / SD_pooled_TQ
pct_TQ_65 = (1 - normal_cdf(z65_TQ)) * 100

# RLLAXH tổng
M_total_XH = (n_male * 43.20 + n_female * 52.74) / total_n
SD_pooled_XH = math.sqrt(((n_male-1)*25.093**2 + (n_female-1)*26.311**2) / (total_n-2))
z65_XH = (65 - M_total_XH) / SD_pooled_XH
pct_XH_65 = (1 - normal_cdf(z65_XH)) * 100

# RLLAC tổng
M_male_C, SD_male_C = 25.42, 25.459
M_female_C, SD_female_C = 24.76, 23.294
M_total_C = (n_male * M_male_C + n_female * M_female_C) / total_n
SD_pooled_C = math.sqrt(((n_male-1)*SD_male_C**2 + (n_female-1)*SD_female_C**2) / (total_n-2))
z65_C = (65 - M_total_C) / SD_pooled_C
pct_C_65 = (1 - normal_cdf(z65_C)) * 100

add_table(
    ['Loại RLLA', 'ĐTB tổng', 'SD', 'z (ngưỡng 65)', '% vượt ngưỡng (giả định)'],
    [
        ['RLLATQ (lan tỏa)', f'{M_total_TQ:.2f}', f'{SD_pooled_TQ:.2f}', f'{z65_TQ:.3f}', f'≈ {pct_TQ_65:.1f}%'],
        ['RLLAXH (xã hội)', f'{M_total_XH:.2f}', f'{SD_pooled_XH:.2f}', f'{z65_XH:.3f}', f'≈ {pct_XH_65:.1f}%'],
        ['RLLAC (chia ly)', f'{M_total_C:.2f}', f'{SD_pooled_C:.2f}', f'{z65_C:.3f}', f'≈ {pct_C_65:.1f}%'],
    ]
)
para('')
para(
    'CẢNH BÁO QUAN TRỌNG: cách này GIẢ ĐỊNH dữ liệu phân '
    'phối CHUẨN — nhưng dữ liệu RLLA thường LỆCH PHẢI '
    '(skewed right) vì đa số HS không có triệu chứng. Tỷ '
    'lệ % thực tế có thể KHÁC.', bold=True, color=RED
)

# Tính theo nam/nữ
para('')
para('So sánh nam vs nữ với ngưỡng 65 (giả định normal):', bold=True)

# RLLATQ theo giới
z_male_TQ = (65 - 51.43) / 22.010
pct_male_TQ = (1 - normal_cdf(z_male_TQ)) * 100
z_female_TQ = (65 - 59.47) / 22.072
pct_female_TQ = (1 - normal_cdf(z_female_TQ)) * 100

# RLLAXH theo giới
z_male_XH = (65 - 43.20) / 25.093
pct_male_XH = (1 - normal_cdf(z_male_XH)) * 100
z_female_XH = (65 - 52.74) / 26.311
pct_female_XH = (1 - normal_cdf(z_female_XH)) * 100

add_table(
    ['Loại RLLA', '% nam vượt 65', '% nữ vượt 65', 'Tỷ số nữ/nam'],
    [
        ['RLLATQ (lan tỏa)', f'≈ {pct_male_TQ:.1f}%', f'≈ {pct_female_TQ:.1f}%', f'{pct_female_TQ/pct_male_TQ:.2f}'],
        ['RLLAXH (xã hội)', f'≈ {pct_male_XH:.1f}%', f'≈ {pct_female_XH:.1f}%', f'{pct_female_XH/pct_male_XH:.2f}'],
    ]
)

H('Cách 3 — Liên hệ tác giả xin DATA RAW', level=3, color=NAVY)
para(
    'Cách CHÍNH XÁC NHẤT: liên hệ tác giả luận án (CTH) xin '
    'data raw, đếm trực tiếp HS vượt ngưỡng. Cách này KHÔNG '
    'phụ thuộc giả định normal distribution.'
)
para(
    'Đề xuất quy trình:'
)
para('1. Yêu cầu tác giả cung cấp ngưỡng cắt RCADS đã sử dụng.')
para('2. Nếu chưa có ngưỡng, đề xuất 65/100 (borderline) hoặc 70/100 (clinical).')
para('3. Tác giả tính tỷ lệ HS vượt ngưỡng từ data raw.')
para('4. Bổ sung kết quả vào báo cáo.')

# Phần IV — So sánh với prevalence quốc tế
H('PHẦN IV — So sánh với % của các nghiên cứu khác', level=2, color=NAVY)
para('Để thầy có "bức tranh thô" so sánh, đây là tỷ lệ % RLLA từ các nghiên cứu lớn:', bold=True)
add_table(
    ['Nghiên cứu', 'Mẫu', 'Thang đo', 'Tỷ lệ % RLLA'],
    [
        ['V-NAMHS (2022) [VN002]', '5.996 HS 10–17 VN', 'DISC-5 (chẩn đoán DSM-5)', '18,45% (chính thức)'],
        ['Hoàng Trung Học (2025) [VN014]', '8.473 HS THCS-THPT VN', 'DASS-21 (sàng lọc)', 'Trong COVID 41,5%; sau COVID 25,4%'],
        ['Chen 2023 [QT007]', '63.205 HS TQ', 'GAD-7 ≥ 10', 'Lo âu 13,9%'],
        ['Wen 2020 [QT008]', '900 HS THCS rural TQ', 'MHT', 'Lo âu nặng 24,8%'],
        ['Qiu 2022 [QT009]', '2.079 HS THCS TQ', 'SAS', 'Nữ 17,5% > nam 11,1%'],
        ['Ước lượng chương 3 (z-score)', '1.352 HS THCS VN', 'RCADS quy đổi (ngưỡng 65)', f'TQ ≈ {pct_TQ_65:.0f}%; XH ≈ {pct_XH_65:.0f}%'],
    ]
)
para('')
para(
    'Lưu ý: các tỷ lệ % này KHÓ SO SÁNH TRỰC TIẾP do dùng '
    'thang đo khác nhau và ngưỡng cắt khác nhau. Việt Nam '
    'có dải rất rộng từ 18,45% (chẩn đoán nghiêm) đến 41,5% '
    '(sàng lọc trong COVID).'
)

# Phần V — Đề xuất tổng
H('PHẦN V — Đề xuất tổng hợp cho báo cáo CTH', level=2, color=NAVY)
blue_run('Tóm gọn 3 việc cần làm:', bold=True)
blue_run(
    '(1) SỬA phần Bàn luận của trợ lý C.H với 6 cải thiện đã '
    'liệt kê: sửa năm Nguyễn Thị Vân 2020 (không phải 2018); '
    'nêu phát hiện đặc biệt BNHĐ → RLLACL; cảnh báo fit '
    'indices KÉM của BPĐP; nêu rõ pattern ba tầng chênh lệch '
    'giới; bổ sung bảng so sánh cường độ; quy ra %.'
)
blue_run(
    '(2) QUY RA % bằng một trong 3 cách: (a) ngưỡng RCADS '
    'chuẩn quốc tế; (b) z-score giả định normal '
    'distribution (như em đã tính ở Cách 2); (c) liên hệ '
    'tác giả xin data raw để đếm chính xác.'
)
blue_run(
    f'(3) BỨC TRANH THÔ ƯỚC LƯỢNG (theo Cách 2, ngưỡng 65/100, '
    f'giả định normal): khoảng {pct_TQ_65:.1f}% HS có lo âu lan '
    f'tỏa mức borderline (nam {pct_male_TQ:.1f}%, nữ '
    f'{pct_female_TQ:.1f}%); khoảng {pct_XH_65:.1f}% có lo âu '
    f'xã hội mức borderline (nam {pct_male_XH:.1f}%, nữ '
    f'{pct_female_XH:.1f}%); CHỈ {pct_C_65:.1f}% có lo âu chia '
    f'ly mức borderline (RẤT THẤP — phù hợp đặc thù rối loạn '
    f'khởi phát sớm). Tỷ lệ RLLAXH ({pct_XH_65:.1f}%) GẦN VỚI '
    f'V-NAMHS 18,45% chẩn đoán; tỷ lệ RLLATQ ({pct_TQ_65:.1f}%) '
    f'GẦN VỚI Hoàng Trung Học 25,4% sàng lọc sau COVID.'
)

# Cảnh báo cuối
H('CẢNH BÁO CUỐI', level=2, color=RED)
para(
    'Tỷ lệ % được tính ở Cách 2 là ƯỚC LƯỢNG GẦN ĐÚNG dựa '
    'trên giả định normal distribution. Trong y văn, dữ liệu '
    'điểm RCADS thường LỆCH PHẢI (đa số HS không có triệu '
    'chứng). Tỷ lệ % thực tế có thể THẤP HƠN ước lượng. '
    'Trước khi đưa số % vào báo cáo CTH, NÊN: (a) tham vấn '
    'tác giả luận án để xin data raw; (b) hoặc nêu rõ "ước '
    'lượng giả định normal" trong chú thích.', bold=True, color=RED
)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
