# -*- coding: utf-8 -*-
"""Dịch đầy đủ QT29 + QT30 + QT31"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save_and_report(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables, {len([x for x in d.paragraphs if x.text.strip()])} paras')

# =====================================================================
# QT29 — BMC Psychiatry NMA 2025 — Li et al.
# =====================================================================
def dich_QT29():
    print('QT29 BMC NMA...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1186/s12888-025-07227-y', size=10)
    add_heading(doc, 'Hiệu quả các loại can thiệp khác nhau cho rối loạn lo âu ở trẻ em và thanh thiếu niên: Tổng quan hệ thống và phân tích tổng hợp mạng Bayesian', 1)
    h = doc.add_paragraph(); r = h.add_run('Effects of Different Interventions on Anxiety Disorders in Children and Adolescents: A Systematic Review and Bayesian Network Meta-Analysis'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Effects of Different Interventions on Anxiety Disorders in Children and Adolescents: A Systematic Review and Bayesian Network Meta-Analysis'),
        ('Tác giả', 'Longhui Li, Qiner Li, Jingyi Wang, Quan Fu, Meng Chi'),
        ('Cơ quan', 'Capital University of Physical Education and Sports, Beijing; Shenyang Pharmaceutical University'),
        ('Tạp chí', 'BMC Psychiatry (Q1, IF ≈ 4,4)'),
        ('Xuất bản', '2025, Vol. 25, Article 809, 14 trang'),
        ('DOI', '10.1186/s12888-025-07227-y'),
        ('Loại NC', 'Tổng quan hệ thống + phân tích tổng hợp mạng Bayesian (NMA)'),
        ('Mẫu', '30 RCT, 1.711 người tham gia, 12 quốc gia'),
        ('Đăng ký', 'PROSPERO: CRD42024587910'),
    ])
    add_page_ref(doc, '1–14', 'BMC Psychiatry', '2025, 25:809')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Bối cảnh: Rối loạn lo âu phổ biến nhất ở trẻ em/VTN: 4,4% trẻ 10–14 tuổi, 5,5% trẻ 15–19 tuổi (WHO). Tỷ lệ tăng sau COVID. Tỷ lệ thuyên giảm sau điều trị chỉ ~50%. Cần xác định can thiệp hiệu quả nhất.')
    p = doc.add_paragraph()
    r = p.add_run('Phương pháp: Tìm kiếm hệ thống 5 cơ sở dữ liệu (Cochrane, Embase, PubMed, Scopus, Web of Science). NMA Bayesian bằng R Studio. Đánh giá chất lượng: GRADE. Từ 19.442 bài → 30 RCT, 1.711 người tham gia. 4 can thiệp: ACT, CBT, PE, VRET vs đối chứng.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: ACT hiệu quả nhất (MD = −3,83; CrI: −9,33 đến 1,51; SUCRA = 0,69). CBT xếp 2 (MD = −3,64; CrI: −7,36 đến −0,48; SUCRA = 0,66) — DUY NHẤT có CrI không bao gồm 0. VRET xếp 3 (SUCRA = 0,51). PE xếp 4 (SUCRA = 0,51). Tất cả 4 can thiệp hiệu quả hơn đối chứng.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['BMC Psychiatry Q1. NMA Bayesian — phương pháp mạnh, so sánh gián tiếp.','30 RCT, 1.711 VTN, 12 quốc gia. PROSPERO tiền đăng ký.','ACT xếp hạng 1 (BẤT NGỜ — CBT thường được coi hàng đầu). Tuy nhiên CrI bao gồm 0.','CBT có BẰNG CHỨNG MẠNH NHẤT (CrI KHÔNG bao gồm 0; 16/30 RCT).','PE (hoạt động thể chất) cũng hiệu quả — DỄ triển khai tại trường VN, chi phí thấp.','VRET hiệu quả — hướng tương lai.','Consistency test + Egger test OK → kết quả đáng tin cậy.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['GRADE: bằng chứng THẤP (low certainty) — imprecision lớn, CrI rộng.','ACT chỉ 6 RCT, PE chỉ 2 RCT — quá ít để kết luận chắc chắn.','Hầu hết RCT từ Mỹ (7), Úc (5), Đức (3) — ít châu Á. 0 RCT từ VN.','Dị chất đáng kể (τ = 5,87).','Tuổi TB 12,4 — chủ yếu trẻ em, ít VTN lớn (15–17).','Không phân biệt loại lo âu (GAD vs SAD vs phobia).']:
        add_p(doc, f'• {b}')

    add_page_ref(doc, '1–5', 'BMC Psychiatry', '2025')
    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Rối loạn lo âu phổ biến nhất ở trẻ em/VTN: 4,4% trẻ 10–14, 5,5% trẻ 15–19 (WHO). Bao gồm: phobia cụ thể, agoraphobia, lo âu xã hội, GAD, rối loạn hoảng sợ, lo âu chia ly. Thường phát triển trước 14 tuổi — liên quan giảm chất lượng sống, bệnh đi kèm, bỏ học, tự tử. COVID-19 làm tăng tỷ lệ do cô lập xã hội, đóng cửa trường, giảm hoạt động thể chất.')
    add_p(doc, 'Tỷ lệ thuyên giảm sau điều trị chỉ ~50% → cần xác định can thiệp hiệu quả nhất. Các can thiệp chính:')
    add_p(doc, '• ACT (Acceptance and Commitment Therapy): Thế hệ 3 CBT. Không nhắm vào THAY ĐỔI suy nghĩ tiêu cực (như CBT truyền thống) mà CHẤP NHẬN chúng + hành động theo giá trị cá nhân. Dựa trên 6 quy trình: chấp nhận, defusion, bản ngã quan sát, hiện tại, giá trị, hành động cam kết.')
    add_p(doc, '• CBT (Cognitive-Behavioral Therapy): Tiêu chuẩn vàng. Tái cấu trúc nhận thức + phơi nhiễm dần. Cochrane review (James et al. 2020): hiệu quả NNT ≈ 3.')
    add_p(doc, '• PE (Physical Exercise): Cơ chế: tăng β-endorphin (cải thiện tâm trạng), điều hòa trục HPA (ổn định cortisol), tăng BDNF (bảo vệ thần kinh). Mức BDNF thường GIẢM ở người lo âu → PE có thể khôi phục.')
    add_p(doc, '• VRET (Virtual Reality Exposure Therapy): Phơi nhiễm qua môi trường ảo. Nguyên lý: kích hoạt mạng sợ hãi trong não → dần giảm nhạy qua habituation. Hiệu quả giảm đau và lo âu ở VTN trải qua thủ thuật y tế.')
    add_p(doc, 'NMA Bayesian: phương pháp so sánh NHIỀU can thiệp đồng thời, kể cả khi chưa so sánh trực tiếp. Tích hợp thông tin tiền nghiệm. Phù hợp mẫu nhỏ.')

    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Tìm kiếm hệ thống: Cochrane, Embase, PubMed, Scopus, Web of Science. Tiêu chuẩn: RCT, trẻ em/VTN có lo âu, so sánh ACT/CBT/PE/VRET vs đối chứng (danh sách chờ, placebo, hoặc TAU). PRISMA + PROSPERO (CRD42024587910). Kappa inter-rater: 30% bài full-text được 2 người đánh giá độc lập.')
    add_p(doc, 'Mô hình: Bayesian random effects NMA (R Studio, gemtc package). DIC (Deviance Information Criterion) chọn mô hình. SUCRA xếp hạng. Consistency test: DIC consistent vs inconsistent. Node-splitting: kiểm tra inconsistency cục bộ. Publication bias: funnel plot + Egger test.')
    add_p(doc, 'Chất lượng: Cochrane Risk of Bias (6 chỉ số) + GRADE (đánh giá tổng thể bằng chứng).')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_p(doc, '3.1. Đặc điểm NC', bold=True)
    add_p(doc, '19.442 bài → 30 RCT, 1.711 người tham gia. 12 quốc gia: Mỹ (7), Úc (5), Đức (3), UK (2), Iran (2), Thụy Điển (2), khác (9). Tuổi TB = 12,40. 62% nữ. Can thiệp: ACT (6), CBT (16), PE (2), VRET (6). Thời gian can thiệp trung vị: 10 tuần, 1 lần/tuần, 60 phút. Theo dõi trung vị: 3 tháng.')
    add_p(doc, 'Thang đo lo âu phổ biến nhất: SCAS (Spence Children\'s Anxiety Scale, 7 NC), MASC (Multidimensional Anxiety Scale for Children, 5 NC), PSAS (Public Speaking Anxiety Scale, 2 NC).')

    add_heading(doc, 'Bảng 1. Xếp hạng hiệu quả can thiệp (NMA Bayesian)', 3)
    add_table(doc,
        ['Can thiệp', 'Số RCT', 'MD vs đối chứng', 'KTC 95% CrI', 'SUCRA', 'Xếp hạng TB', 'Kết luận'],
        [['ACT', '6', '−3,83', '−9,33 đến 1,51', '0,69', '2,25', 'Hạng 1 nhưng CrI qua 0'],
         ['CBT', '16', '−3,64', '−7,36 đến −0,48', '0,66', '2,31', 'Hạng 2 — BẰNG CHỨNG MẠNH NHẤT'],
         ['VRET', '6', '−2,53', '−8,23 đến 3,32', '0,51', '2,86', 'Hạng 3 — CrI rộng'],
         ['PE', '2', '−2,16', '−9,99 đến 5,52', '0,51', '3,12', 'Hạng 4 — chỉ 2 RCT'],
         ['Đối chứng', '—', 'Ref', '—', '—', '—', '']],
        widths=[2.0, 1.0, 2.0, 3.0, 1.5, 1.5, 3.5])
    add_p(doc, 'CBT là can thiệp DUY NHẤT có CrI không bao gồm 0 — bằng chứng chắc chắn nhất. ACT xếp hạng 1 (SUCRA 0,69) nhưng CrI bao gồm 0 (−9,33 đến 1,51) → bất định.', size=9, italic=True)

    add_p(doc, '3.2. Chất lượng bằng chứng', bold=True)
    add_p(doc, 'Risk of Bias: phân bổ ngẫu nhiên rõ ràng (22/30 NC thấp); ẩn phân bổ thấp (20/30); mù (hầu hết không mù — khó khả thi cho can thiệp tâm lý); dữ liệu đầy đủ (23/30). 1 NC có bias báo cáo chọn lọc. 11 NC không có bias khác.')
    add_p(doc, 'GRADE: bằng chứng tổng thể ở mức THẤP — do imprecision (CrI rộng) và thiên lệch xuất bản tiềm năng. Mặc dù Egger test p = 0,91 (không bias), mẫu nhỏ cho một số so sánh.')
    add_p(doc, 'Consistency: DIC consistent (141,3) < inconsistent (142,6) → nhất quán. Node-splitting: tất cả p > 0,05 → không inconsistency cục bộ. Dị chất: τ = 5,87 (95% CrI: 2,92–9,83) — ĐÁNG KỂ, phản ánh khác biệt thiết kế NC.')

    add_heading(doc, 'Bảng 2. Giải thích các loại can thiệp', 3)
    add_table(doc,
        ['Can thiệp', 'Tên đầy đủ', 'Cơ chế', 'Ưu/nhược điểm cho VN'],
        [['ACT', 'Acceptance and Commitment Therapy', 'Chấp nhận cảm xúc + hành động theo giá trị. 6 quy trình.', 'MỚI ở VN; cần đào tạo chuyên gia; tiềm năng văn hóa phù hợp (Phật giáo/chấp nhận)'],
         ['CBT', 'Cognitive-Behavioral Therapy', 'Tái cấu trúc nhận thức + phơi nhiễm dần (inhibitory learning).', 'Phổ biến nhất; CBT nhóm khả thi; CAMS 80,7% kết hợp SSRI (AJP QT28)'],
         ['PE', 'Physical Exercise', 'Endorphin ↑, BDNF ↑, HPA ổn định, cortisol ↓.', 'DỄ TRIỂN KHAI NHẤT tại trường VN; chi phí thấp; không cần chuyên gia tâm lý'],
         ['VRET', 'Virtual Reality Exposure Therapy', 'Phơi nhiễm qua VR → habituation → giảm nhạy sợ hãi.', 'Cần công nghệ + chi phí; tiềm năng tương lai']],
        widths=[2.0, 3.5, 4.5, 4.5])

    add_heading(doc, '4. THẢO LUẬN', 2)
    add_p(doc, 'NC xác nhận tất cả 4 can thiệp hiệu quả cho lo âu trẻ em/VTN. ACT xếp hạng 1 (SUCRA 0,69) — BẤT NGỜ vì CBT thường được coi là tiêu chuẩn vàng. Tuy nhiên, ACT chỉ có 6 RCT và CrI bao gồm 0 → cần thận trọng. CBT có bằng chứng mạnh nhất (16 RCT, CrI không qua 0).')
    add_p(doc, 'ACT: Thế hệ 3 CBT, nhấn mạnh CHẤP NHẬN thay vì THAY ĐỔI suy nghĩ. Có thể phù hợp VTN không đáp ứng CBT truyền thống. Cũng có thể phù hợp văn hóa châu Á (Phật giáo, triết lý chấp nhận).')
    add_p(doc, 'PE: Đáng chú ý — ít NMA nào bao gồm. Cơ chế sinh học rõ ràng (endorphin, BDNF, HPA). DỄ triển khai tại trường (thể dục, yoga). Tuy nhiên chỉ 2 RCT → cần thêm NC.')
    add_p(doc, 'VRET: Hiệu quả cho phobia cụ thể và lo âu xã hội. Giảm đau/lo âu ở VTN trải qua thủ thuật y tế. Cần thêm RCT với mẫu lớn.')
    add_p(doc, 'So sánh với AJP 2024 (QT28, Zugman et al.): AJP tập trung CBT + SSRIs + cơ chế thần kinh. KHÔNG đề cập ACT hoặc PE — thiếu sót quan trọng mà NMA này bổ sung.')

    add_heading(doc, '5. KẾT LUẬN', 2)
    add_p(doc, 'NMA 30 RCT (1.711 VTN, 12 quốc gia) xác nhận ACT, CBT, PE, VRET đều hiệu quả. ACT xếp hạng 1 (SUCRA 0,69) nhưng bằng chứng chưa chắc chắn. CBT có bằng chứng mạnh nhất. PE hiệu quả và dễ triển khai. Bằng chứng tổng thể THẤP → cần thêm RCT chất lượng cao, đặc biệt ở LMIC/châu Á/VN.')

    add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
    for ref in ['Li, L.H., et al. (2025). Effects of different interventions on anxiety disorders in children and adolescents: NMA. BMC Psychiatry, 25, 809.','James, A.C., et al. (2020). CBT for anxiety in children. Cochrane Database Syst Rev, 11, CD004690.','Zugman, A., et al. (2024). Current and future approaches to pediatric anxiety treatment. AJP, 181, 189–200.','Walkup, J.T., et al. (2008). CBT, sertraline, or combination in childhood anxiety (CAMS). NEJM, 359, 2753–2766.','(Xem đầy đủ trong bài gốc — 107 TLTK)']:
        add_p(doc, ref, size=10)

    add_abbreviation_table(doc, [('NMA','Network Meta-Analysis'),('ACT','Acceptance and Commitment Therapy'),('CBT','Cognitive-Behavioral Therapy'),('PE','Physical Exercise'),('VRET','Virtual Reality Exposure Therapy'),('SUCRA','Surface Under the Cumulative Ranking'),('MD','Mean Difference'),('CrI','Credible Interval (Bayesian)'),('DIC','Deviance Information Criterion'),('GRADE','Grading of Recommendations Assessment'),('BDNF','Brain-Derived Neurotrophic Factor'),('HPA','Hypothalamic-Pituitary-Adrenal'),('SCAS','Spence Children\'s Anxiety Scale'),('MASC','Multidimensional Anxiety Scale for Children'),('PROSPERO','International Prospective Register of Systematic Reviews')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['BMC Psychiatry Q1. NMA Bayesian — phương pháp mạnh nhất cho so sánh gián tiếp. PROSPERO tiền đăng ký.','30 RCT, 12 quốc gia — bao quát. Consistency + Egger OK.','Bao gồm ACT + PE — ít NMA nào làm. AJP 2024 (QT28) KHÔNG đề cập ACT/PE.','ACT hạng 1 — phát hiện mới, mở rộng lựa chọn ngoài CBT.','PE — cơ chế sinh học rõ (endorphin, BDNF, HPA). Dễ triển khai tại trường VN (Zhu 2025 QT05: hoạt động ngoài trời bảo vệ SKTT).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế chi tiết:', bold=True)
    for s in ['GRADE THẤP — imprecision lớn (CrI rộng, bao gồm 0 cho ACT/PE/VRET). CHỈ CBT có CrI không qua 0.','ACT chỉ 6 RCT — THIẾU để kết luận chắc chắn. PE chỉ 2 RCT — rất ít.','Hầu hết RCT từ phương Tây (Mỹ 7, Úc 5, Đức 3) — 0 RCT từ VN/ĐNA. Zhameden 2025 (QT03): 0 RCT VN.','Dị chất τ = 5,87 — đáng kể. NC rất khác nhau: tuổi (4–18), loại lo âu, thời gian, đo lường.','Tuổi TB 12,4 — chủ yếu trẻ em (pre-teen), ít VTN lớn (15–17 tuổi — nhóm lo âu tăng nhanh nhất theo JAACAP QT23: AOR = 2,93 ở 15–17).','Không phân biệt loại lo âu — GAD, SAD, phobia gộp chung. AJP 2024 (QT28): 3 loại riêng. Jefferies 2020 (QT35): lo âu XH riêng.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN 0 RCT can thiệp lo âu VTN — GAP #1 cross-study. CBT nhóm + PE tại trường VN: ưu tiên cao nhất.','PE tại trường VN — DỄ NHẤT, RẺ NHẤT. Thể dục + yoga + vận động ngoài trời. BMC NMA: SUCRA 0,51.','ACT tại VN — mới nhưng tiềm năng. Cần đào tạo. Thảo Vi 2025 (VN19): lạc quan (gần ACT) β gián tiếp = −0,24.','RCT ở VTN 15–17 tuổi — nhóm tăng nhanh nhất (JAACAP QT23) nhưng ÍT RCT nhất.','Kết hợp can thiệp: CBT + PE + giảm screen time — chưa có RCT nào thử. JAMA 2024 (QT33): giảm ST d=0,53.']:
        add_red(doc, f'• {s}')
    save_and_report(doc, '29_CBT_NetworkMeta_2025_BMCPsych.docx')

# =====================================================================
# QT30 — GBD Trends 2025 — Zhang et al.
# =====================================================================
def dich_QT30():
    print('QT30 GBD Trends...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.04.001', size=10)
    add_heading(doc, 'Xu hướng rối loạn trầm cảm và lo âu ở thanh thiếu niên và thanh niên (10–24 tuổi) từ 1990 đến 2021: Phân tích Gánh nặng Bệnh tật Toàn cầu', 1)
    h = doc.add_paragraph(); r = h.add_run('Trends in Depressive and Anxiety Disorders among Adolescents and Young Adults (aged 10–24) from 1990 to 2021: A Global Burden of Disease Study Analysis'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Zhang Dongjun, Wu Mingyue, Li Xinqi, Wang Lina, Wu Jiali, Jin Mengyao'),
        ('Cơ quan', 'Xinxiang Medical University, Trung Quốc'),
        ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
        ('Xuất bản', '2025, 14 trang'),
        ('DOI', '10.1016/j.jad.2025.04.001'),
        ('Loại NC', 'Phân tích xu hướng sinh thái hồi cứu — GBD 2021, Joinpoint regression'),
        ('Phạm vi', '204 quốc gia, 1990–2021, VTN và thanh niên 10–24 tuổi'),
    ])
    add_page_ref(doc, '1–14', 'J Affective Disorders', '2025')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Mục tiêu: Đánh giá xu hướng dài hạn gánh nặng trầm cảm và lo âu ở 10–24 tuổi từ 1990 đến 2021. Phương pháp: Joinpoint regression — AAPC cho tỷ lệ mới mắc và DALYs. Bất bình đẳng: SII + CI. Phân tầng: tuổi, giới, khu vực, quốc gia, SDI.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: AAPC toàn cầu: trầm cảm 0,97% (KTC: 0,75–1,19), lo âu 0,84% (0,48–1,21). Tăng tốc 2014–2021, đỉnh 2019. Trầm cảm: 3.085,5 → 3.909,9/100K. Lo âu: 708,0 → 883,1/100K. Nữ gánh nặng cao hơn. 10–14 tuổi AAPC cao nhất (1,44%). SDI cao tăng nhanh nhất. Tổng: 73,8 triệu ca mới trầm cảm + lo âu ở 10–24 tuổi năm 2021.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['J Affective Disorders Q1. GBD 2021 — 204 quốc gia, 31 năm, toàn diện nhất.','Joinpoint regression — xác định CHÍNH XÁC điểm thay đổi xu hướng (2006, 2010, 2014, 2019). Tăng tốc từ 2014.','AAPC: trầm cảm 0,97%/năm, lo âu 0,84%/năm — tăng liên tục.','10–14 tuổi AAPC CAO NHẤT (1,44%) — rối loạn khởi phát ngày càng SỚM.','Phân tích bất bình đẳng (SII, CI) — ít NC nào làm. SDI cao tăng nhanh nhất.','73,8 triệu ca mới trầm cảm + lo âu ở 10–24 tuổi năm 2021 toàn cầu.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['GBD dựa MÔ HÌNH ước tính — nhiều nước (VN) thiếu dữ liệu gốc.','Sinh thái hồi cứu — không nhân quả.','Không tách loại lo âu.','SDI thấp tăng chậm — có thể do thiếu dữ liệu, không phải thực sự thấp.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Trầm cảm: nguyên nhân thứ 2 gây YLDs toàn cầu. Lo âu: thứ 6 (Lancet Psychiatry, 2024). Tuy nhiên, NC xu hướng ở VTN 10–24 tuổi hạn chế, đặc biệt về mối quan hệ với SDI.')
    add_p(doc, '4 câu hỏi NC: (1) Xu hướng theo tuổi, giới, khu vực, SDI; (2) Bất bình đẳng giữa các nước 1990–2021; (3) Bất bình đẳng xuyên quốc gia; (4) Bằng chứng cho chính sách can thiệp.')

    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'GBD 2021: 371 bệnh/thương tích, 204 quốc gia, 100.983 đầu vào dữ liệu. VTN 10–24 tuổi: 3 nhóm (10–14, 15–19, 20–24). Trầm cảm + Lo âu theo DSM-IV/ICD-10.')
    add_p(doc, 'Phân tích: Joinpoint regression — xác định điểm thay đổi xu hướng (joinpoints). AAPC cho toàn giai đoạn. APC cho từng đoạn. SII (Slope Index of Inequality) + CI (Concentration Index) cho bất bình đẳng theo SDI.')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xu hướng toàn cầu 1990–2021', 3)
    add_table(doc,
        ['Rối loạn / Phân nhóm', 'Mới mắc 1990 (/100K)', 'Mới mắc 2021 (/100K)', 'AAPC (%)', 'p'],
        [['Trầm cảm — tổng', '3.085,5 (2.307,8–4.105,6)', '3.909,9 (2.841,3–5.171,1)', '0,97 (0,75–1,19)', '<0,001'],
         ['Trầm cảm — nam', '2.265,2', '2.991,7', '1,07', '<0,001'],
         ['Trầm cảm — nữ', '3.932,7', '4.875,1', '0,93', '<0,001'],
         ['Trầm cảm — 10–14', '—', '—', '1,44', '<0,001'],
         ['Lo âu — tổng', '708,0 (511,5–939,3)', '883,1 (636,9–1.167,1)', '0,84 (0,48–1,21)', '<0,001'],
         ['Lo âu — 10–14', '397,1', '547,2', '1,12', '<0,001'],
         ['Lo âu — 15–19', '943,6', '1.180,8', '0,93', '<0,001'],
         ['Lo âu — 20–24', '1.105,8', '1.287,1', '0,60', '<0,001']],
        widths=[3.5, 3.0, 3.0, 2.5, 1.5])
    add_p(doc, 'AAPC = thay đổi %/năm trung bình. Nhóm 10–14 tăng NHANH NHẤT (1,44% trầm cảm, 1,12% lo âu).', size=9, italic=True)

    add_heading(doc, 'Bảng 2. Xu hướng theo khu vực (AAPC cao nhất)', 3)
    add_table(doc,
        ['Khu vực', 'AAPC Trầm cảm', 'AAPC Lo âu', 'Ghi chú'],
        [['Bắc Mỹ thu nhập cao', '2,30 (1,53–3,07)', '—', 'Cao nhất trầm cảm'],
         ['Mỹ Latinh Trung', '2,00 (1,72–2,28)', '1,48 (1,18–1,79)', ''],
         ['Mỹ Latinh Andes', '1,68 (1,22–2,14)', '1,91 (1,82–2,00)', 'Cao nhất lo âu'],
         ['Mỹ Latinh nhiệt đới', '—', '1,61 (1,42–1,81)', ''],
         ['SDI cao', 'Cao nhất', 'Cao nhất', 'Tương quan SDI dương'],
         ['SDI thấp', 'Thấp nhất', 'Thấp nhất', 'Có thể do thiếu dữ liệu']],
        widths=[4.0, 3.0, 3.0, 3.5])

    add_p(doc, '3.2. Joinpoint — Điểm thay đổi xu hướng', bold=True)
    add_p(doc, 'Trầm cảm: Joinpoints tại 2006, 2010, 2019. Tăng 1990–1997, 1998–2005, 2014–2021. Giảm 2006–2013. ĐỈNH 2019.')
    add_p(doc, 'Lo âu: Tăng 2006–2013 và 2014–2021. Giảm 1998–2005. Ổn định 1990–1997.')
    add_p(doc, 'Cả hai: TĂNG TỐC từ 2014 — trùng với bùng nổ smartphone/MXH toàn cầu (phù hợp Norway 2025 QT21: MXH giải thích xu hướng).')

    add_p(doc, '3.3. Bất bình đẳng', bold=True)
    add_p(doc, 'SII dương → bất bình đẳng theo SDI: nước SDI cao có gánh nặng CAO HƠN. Paradox: nước giàu hơn → nhiều lo âu/trầm cảm hơn? Có thể do: (1) phát hiện tốt hơn, (2) MXH/screen time nhiều hơn, (3) áp lực thành tích cao hơn. Korea 2024 (QT34): stress nhóm giàu 40,07% — vẫn rất cao.')

    add_heading(doc, '4. THẢO LUẬN VÀ KẾT LUẬN', 2)
    add_p(doc, 'GBD 204 quốc gia 31 năm xác nhận: trầm cảm và lo âu ở 10–24 tuổi TĂNG liên tục toàn cầu. Tăng tốc từ 2014, đỉnh 2019. Nữ gánh nặng cao hơn. 10–14 tuổi tăng nhanh nhất — rối loạn khởi phát ngày càng sớm. Bất bình đẳng: SDI cao tăng nhanh nhất — paradox nước giàu.')
    add_p(doc, 'Hàm ý: Cần can thiệp SỚM (nhắm 10–14 tuổi — THCS), đặc biệt cho NỮ. VN nằm trong SDI trung bình — ước tính GBD có thể không chính xác (thiếu dữ liệu gốc). Cần khảo sát quốc gia lặp lại (V-NAMHS 2022 là thời điểm 1).')

    add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
    for ref in ['Zhang, D., et al. (2025). Trends in depressive and anxiety disorders among 10–24 year olds 1990–2021: GBD analysis. J Affective Disorders.','GBD 2021 Study (2024). Global burden of 371 diseases. Lancet.','(Xem đầy đủ trong bài gốc — 50+ TLTK)']:
        add_p(doc, ref, size=10)

    add_abbreviation_table(doc, [('GBD','Global Burden of Disease'),('AAPC','Average Annual Percent Change'),('APC','Annual Percent Change'),('DALYs','Disability-Adjusted Life Years'),('YLDs','Years Lived with Disability'),('SDI','Socio-Demographic Index'),('SII','Slope Index of Inequality'),('CI','Concentration Index'),('UI','Uncertainty Interval'),('DSM-IV','Diagnostic and Statistical Manual 4th Edition'),('ICD-10','International Classification of Diseases 10th Revision')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['J Affective Disorders Q1. GBD 2021 — DỮ LIỆU TOÀN DIỆN NHẤT (204 nước, 31 năm, 100K+ đầu vào).','Joinpoint regression — xác định CHÍNH XÁC năm thay đổi (2006, 2010, 2014, 2019). 2014 tăng tốc → trùng smartphone/MXH bùng nổ. Phù hợp Norway 2025 (QT21).','AAPC cho từng phân nhóm — so sánh chi tiết tuổi × giới × khu vực × SDI.','Bất bình đẳng (SII, CI) — đóng góp mới. Korea 2024 (QT34) phân tầng thu nhập; bài này phân tầng SDI quốc gia.','10–14 tuổi tăng NHANH NHẤT — phát hiện quan trọng. VN cần NC riêng cho THCS (10–14).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['GBD dựa MÔ HÌNH ước tính — nhiều nước (VN) thiếu dữ liệu khảo sát quốc gia → ước tính có thể KHÔNG CHÍNH XÁC. V-NAMHS 2022 là một trong ít nguồn gốc VN.','Sinh thái hồi cứu — xu hướng cấp QUỐC GIA, không cá nhân. Không xác lập nhân quả.','Không tách loại lo âu — GAD, SAD, phobia gộp chung. JAACAP 2024 (QT23) tách riêng: lo âu tăng gấp đôi.','SDI thấp tăng CHẬM — có thể do THIẾU DỮ LIỆU chứ không phải thực sự thấp. Paradox nước giàu cần thận trọng.','Đỉnh 2019 — TRƯỚC COVID. Chưa phản ánh tác động COVID đầy đủ. Korea 2024 (QT34): đảo chiều sau COVID.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN CHƯA CÓ dữ liệu xu hướng dài hạn (>5 năm) lo âu VTN — GBD ước tính chưa xác nhận. Cần khảo sát quốc gia lặp lại mỗi 3–5 năm.','10–14 tuổi (THCS) tăng nhanh nhất — VN CẦN NC riêng cho nhóm này. Hầu hết NC VN tập trung THPT (14–18).','So sánh ước tính GBD với dữ liệu thực VN — V-NAMHS 2022 (DISC-5: 2,3% chẩn đoán). GBD ước tính bao nhiêu cho VN?']:
        add_red(doc, f'• {s}')
    save_and_report(doc, '30_GBD_Trends_10-24y_2025.docx')

# =====================================================================
# QT31 — 59 Countries — Islam et al. 2025
# =====================================================================
def dich_QT31():
    print('QT31 59Countries...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.06.001', size=10)
    add_heading(doc, 'Tỷ lệ và các yếu tố liên quan của lo âu ở thanh thiếu niên đi học: Phân tích từ 59 quốc gia', 1)
    h = doc.add_paragraph(); r = h.add_run('Prevalence of and Factors Associated with Anxiety among School Going Adolescents: Analysis from 59 Countries'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Md. Amirul Islam, Tanjirul Islam, Bristi Rani Saha, Sakib Al Hassan, Noman Hasan, Md. Ashfikur Rahman'),
        ('Cơ quan', "Xi'an Jiao Tong University; Khulna University, Bangladesh; City University of Hong Kong"),
        ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
        ('Xuất bản', '2025, Vol. 393, 120315, 11 trang'),
        ('DOI', '10.1016/j.jad.2025.06.001'),
        ('Loại NC', 'Phân tích thứ cấp — GSHS (Global School-based Student Health Survey, WHO/CDC)'),
        ('Mẫu', '179.937 VTN 11–17 tuổi đi học, từ 59 quốc gia (chủ yếu LMIC)'),
    ])
    add_page_ref(doc, '1–11', 'J Affective Disorders', '2025')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Bối cảnh: Tuổi VTN là giai đoạn quan trọng với thay đổi cảm xúc/tâm lý đáng kể. NC nhằm điều tra tỷ lệ lo âu và yếu tố liên quan ở VTN đi học trên 59 quốc gia LMIC.')
    p = doc.add_paragraph()
    r = p.add_run('Phương pháp: Phân tích dữ liệu GSHS (WHO/CDC), 179.937 VTN 11–17 tuổi từ 59 quốc gia. Lo âu đánh giá dựa trên câu hỏi GSHS chuẩn hóa. Hồi quy logistic đa biến (AOR) phân tầng theo 6 khu vực WHO.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Lo âu phổ biến ở VTN đi học toàn cầu. Yếu tố nguy cơ mạnh nhất: ý tưởng tự tử (AOR = 2,84), bất an thực phẩm nặng (AOR = 2,22), tuổi 17 (AOR = 2,45), bắt nạt (AOR = 1,68), chấn thương (AOR = 1,61), ngồi >4h/ngày (AOR = 1,50), nữ (AOR = 1,51). Yếu tố bảo vệ: cha mẹ kiểm tra bài (AOR = 0,75).')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['J Affective Disorders Q1. MẪU CỰC LỚN: 179.937 VTN từ 59 quốc gia LMIC.','GSHS (WHO/CDC) — chuẩn hóa, đại diện, so sánh được giữa các nước.','Bất an thực phẩm (AOR = 2,22) — yếu tố ít được NC trước đó, rất mạnh.','Ý tưởng tự tử AOR = 2,84 — yếu tố mạnh nhất, liên quan lo âu-tự tử mạnh.','Cha mẹ kiểm tra bài AOR = 0,75 — yếu tố bảo vệ. Phù hợp Wen 2020 (OR = 0,562 hỗ trợ trường).','Phân tầng 6 khu vực WHO — phát hiện khác biệt vùng.','Bao gồm nhiều nước ĐNA — gần bối cảnh VN.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Lo âu đo bằng 1 CÂU GSHS — KHÔNG phải thang chuẩn (GAD-7, DASS-21).','Cắt ngang — không nhân quả.','Chỉ VTN ĐI HỌC — bỏ sót bỏ học/thất nghiệp (nhóm nguy cơ cao hơn).','VN KHÔNG nằm trong 59 nước (GSHS VN chưa công bố gần đây).']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Lo âu là rối loạn tâm thần phổ biến nhất ở VTN: 4,4% trẻ 10–14, 5,5% trẻ 15–19 (WHO, 2024). Tỷ lệ cao nhất: Mỹ Latinh, Caribbean, Bắc Mỹ, Tây Âu. Thấp nhất: Nam Á, châu Phi hạ Sahara (Javaid et al., 2023).')
    add_p(doc, 'Lo âu VTN liên quan: suy giảm chức năng, học kém, đau khổ cảm xúc, giảm chất lượng sống (Viswanathan et al., 2022). Tăng nguy cơ: tự trọng thấp, trầm cảm, thuốc lá, nghiện chất, tự tử (Lawrence et al., 2017). Lo âu VTN kéo dài → lo âu người lớn.')
    add_p(doc, 'GSHS: Khảo sát Sức khỏe Học đường Toàn cầu do WHO và CDC phối hợp. Thiết kế: cắt ngang, lấy mẫu cụm 2 bậc, đại diện quốc gia cho VTN đi học 13–17 tuổi. Thực hiện tại >100 quốc gia.')

    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Dữ liệu: GSHS (WHO/CDC), 179.937 VTN 11–17 tuổi từ 59 quốc gia LMIC. Biến phụ thuộc: lo âu — câu hỏi GSHS: "Trong 12 tháng qua, bạn có thường xuyên lo lắng đến mức không thể ngủ không?" (thường xuyên/luôn luôn = có lo âu).')
    add_p(doc, 'Biến độc lập: tuổi, giới, bất an thực phẩm (food insecurity: không/đôi khi/thường xuyên), bắt nạt, đánh nhau, chấn thương, bạn bè thân (0/1–2/≥3), cha mẹ kiểm tra bài/hiểu vấn đề con, ngồi nhiều (<1h/1–4h/>4h/ngày), BMI, ý tưởng/kế hoạch/cố gắng tự tử.')
    add_p(doc, 'Phân tích: Hồi quy logistic đa biến (AOR, KTC 95%). Phân tầng 6 khu vực WHO: Châu Phi, Đông Địa Trung Hải, Đông Nam Á, Châu Mỹ, Tây Thái Bình Dương, Châu Âu.')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Yếu tố liên quan lo âu VTN đi học (59 quốc gia, n = 179.937)', 3)
    add_table(doc,
        ['Yếu tố', 'AOR', 'KTC 95%', 'p', 'Ý nghĩa'],
        [['Tuổi 17 (vs 11–12)', '2,45', '2,05–2,94', '<0,001', 'TĂNG GẤP 2,5 — tuổi lớn → nguy cơ cao'],
         ['Nữ (vs nam)', '1,51', '1,38–1,66', '<0,001', 'Nữ > Nam — nhất quán toàn cầu'],
         ['Bất an thực phẩm — nặng', '2,22', '1,94–2,55', '<0,001', 'YẾU TỐ MẠNH THỨ 2'],
         ['Bắt nạt', '1,68', '1,52–1,85', '<0,001', ''],
         ['Chấn thương nghiêm trọng', '1,61', '1,46–1,79', '<0,001', ''],
         ['Đánh nhau', '1,26', '—', '<0,001', ''],
         ['Không có bạn thân', '1,28', '1,10–1,48', '0,001', ''],
         ['Cha mẹ không hiểu vấn đề con', '1,44', '1,27–1,63', '<0,001', ''],
         ['CHA MẸ KIỂM TRA BÀI', '0,75', '0,67–0,84', '<0,001', 'BẢO VỆ — giảm 25% nguy cơ'],
         ['Ngồi >4h/ngày', '1,50', '1,32–1,71', '<0,001', 'Ít vận động'],
         ['Ý tưởng tự tử', '2,84', '1,85–2,36', '<0,001', 'YẾU TỐ MẠNH NHẤT'],
         ['Kế hoạch tự tử', '1,36', '1,21–1,54', '<0,001', ''],
         ['Cố gắng tự tử', '1,41', '1,25–1,60', '<0,001', '']],
        widths=[4.0, 1.5, 2.5, 1.5, 4.0])

    add_heading(doc, 'Bảng 2. Khác biệt theo khu vực WHO', 3)
    add_table(doc,
        ['Khu vực', 'Đặc điểm nổi bật', 'So sánh'],
        [['Đông Địa Trung Hải', 'Bất an thực phẩm mạnh nhất', 'Xung đột, nghèo'],
         ['Châu Phi', 'Bắt nạt, chấn thương nổi bật', 'Bạo lực'],
         ['Đông Nam Á', 'Tuổi, nữ, thực phẩm mạnh', 'Gần bối cảnh VN'],
         ['Châu Mỹ', 'Ý tưởng tự tử mạnh nhất', ''],
         ['Tây Thái Bình Dương', 'Béo phì nổi bật hơn', 'TQ, Philippines'],
         ['Châu Âu', 'Ít nước tham gia GSHS', '']],
        widths=[3.5, 5.0, 4.0])

    add_heading(doc, '4. THẢO LUẬN', 2)
    add_p(doc, 'Đánh giá TOÀN DIỆN NHẤT về lo âu VTN đi học (59 quốc gia, 179.937 VTN). Phát hiện chính:')
    add_p(doc, '(1) Bất an thực phẩm (AOR = 2,22) — yếu tố ít được NC, rất mạnh. Cơ chế: thiếu dinh dưỡng → stress gia đình → lo âu. Đặc biệt quan trọng ở LMIC.')
    add_p(doc, '(2) Nữ > nam (AOR = 1,51) — nhất quán: bạo lực giới, phân biệt trong trường/xã hội, yếu tố sinh học (estrogen). Phù hợp Ireland (QT32), Úc (QT25), Korea (QT34).')
    add_p(doc, '(3) Cha mẹ kiểm tra bài (AOR = 0,75) — BẢO VỆ quan trọng. Phù hợp Wen 2020 (hỗ trợ trường OR = 0,562), Ireland (QT32, OGA bảo vệ), Korea (QT34: giám sát cha mẹ).')
    add_p(doc, '(4) Ý tưởng tự tử AOR = 2,84 — lo âu và tự tử liên quan RẤT MẠNH. Cần sàng lọc KẾT HỢP lo âu + tự tử, không riêng lẻ. Danh Lâm 2022 (VN17): 23,6% VTN Thanh Hóa đã nghĩ đến tự tử.')

    add_heading(doc, '5. KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu GSHS 179.937 VTN từ 59 quốc gia xác nhận lo âu VTN đi học phổ biến toàn cầu. Yếu tố mạnh nhất: ý tưởng tự tử (2,84), bất an thực phẩm (2,22), tuổi 17 (2,45), bắt nạt (1,68), nữ (1,51). Cha mẹ tham gia (0,75) bảo vệ. Cần can thiệp đa tầng: an ninh lương thực, chống bắt nạt, tăng hỗ trợ gia đình, sàng lọc kết hợp lo âu + tự tử.')

    add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
    for ref in ['Islam, M.A., et al. (2025). Prevalence and factors associated with anxiety among school going adolescents: analysis from 59 countries. J Affective Disorders, 393, 120315.','WHO (2024). Global Health Estimates.','Javaid, S.F., et al. (2023). Epidemiology of anxiety disorders. Frontiers in Psychiatry, 14, 1268354.','(Xem đầy đủ trong bài gốc — 60+ TLTK)']:
        add_p(doc, ref, size=10)

    add_abbreviation_table(doc, [('GSHS','Global School-based Student Health Survey'),('AOR','Adjusted Odds Ratio'),('LMIC','Low and Middle-Income Countries'),('BMI','Body Mass Index'),('WHO','World Health Organization'),('CDC','Centers for Disease Control and Prevention')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['J Affective Disorders Q1. Mẫu CỰC LỚN 179.937 VTN, 59 quốc gia — lớn nhất về lo âu VTN.','GSHS chuẩn hóa WHO/CDC — so sánh được giữa các nước LMIC.','Bất an thực phẩm — đóng góp mới, AOR = 2,22. Ít NC trước đó đánh giá.','Phân tầng 6 khu vực WHO — phát hiện ĐNA riêng biệt.','Cha mẹ bảo vệ (AOR = 0,75) — phù hợp Wen 2020, Ireland QT32 (OGA), Norway QT21 (bất mãn cha mẹ không giải thích).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['Lo âu chỉ 1 CÂU GSHS — KHÔNG phải thang chuẩn. So JAACAP QT23 (chẩn đoán lâm sàng), Norway QT21 (HSCL-6), Ireland QT32 (DASS-21). Tỷ lệ có thể không chính xác.','Cắt ngang — không nhân quả. Bất an thực phẩm → lo âu hay lo âu → ăn kém?','Chỉ VTN ĐI HỌC — bỏ sót bỏ học/thất nghiệp/lang thang (nhóm nguy cơ cao hơn). JAMA QT33: trẻ gia đình có động cơ cao tham gia.','VN KHÔNG nằm trong 59 nước — GSHS VN chưa công bố gần đây. Cần VN tham gia GSHS đầy đủ.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN cần tham gia GSHS đầy đủ + phân tích riêng lo âu — hiện thiếu.','So sánh VN vs ĐNA (Thái Lan, Philippines, Indonesia — có trong GSHS). Jefferies 2020 (QT35): VN 30,7% SAD vs Indonesia 22,9%.','Bất an thực phẩm VN — chưa NC liên quan lo âu VTN, đặc biệt nông thôn/DTTS. Ngô Anh Vinh (VN15): DTTS Lạng Sơn 54,4% — có thể liên quan thực phẩm.','Sàng lọc kết hợp lo âu + tự tử tại trường VN — AOR = 2,84 cho ý tưởng tự tử rất mạnh. Danh Lâm (VN17): 23,6% nghĩ tự tử.']:
        add_red(doc, f'• {s}')
    save_and_report(doc, '31_59Countries_Anxiety_2025.docx')

# RUN ALL
dich_QT29()
dich_QT30()
dich_QT31()
print('\n=== DONE QT29+30+31 ===')
