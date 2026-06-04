# -*- coding: utf-8 -*-
"""Append main body translation to VN002 FULL docx (part 2: Introduction, Sample, Mental Health)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_imgs/'

doc = Document(OUT)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False, underline=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if underline: r.underline = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def H1(text):
    return P(text, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x1F, 0x3A, 0x68))
def H2(text):
    return P(text, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(text):
    return P(text, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))
def H4(text):
    return P(text, bold=True, italic=True, size=11)

def PageMarker(page_num, note='V-NAMHS, UNICEF/IOS 2022'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'--- Trang {page_num}, {note} ---')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

def AddImg(filename, caption_vn, caption_en, width_cm=11.0):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        print(f'  [SKIP] {filename}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(path, width=Cm(width_cm))
    except Exception as e:
        print(f'  [ERR] {filename}: {e}')
        return
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# =============================================================
# EXECUTIVE SUMMARY (page 7)
# =============================================================
PageMarker('7–8')
H2('TÓM TẮT CHO LÃNH ĐẠO (Executive Summary)')

P('Trên toàn cầu, rối loạn sức khoẻ tâm thần thường khởi phát trong giai đoạn vị thành niên, làm đây là giai đoạn then chốt cho can thiệp sớm. Tuy nhiên, tại Việt Nam, dữ liệu đại diện quốc gia về tỷ lệ rối loạn tâm thần ở VTN gần như không có — các nghiên cứu sẵn có chủ yếu dùng thang sàng lọc trên mẫu nhỏ, khiến hoạch định chính sách thiếu cơ sở bằng chứng vững chắc.')

P('Khảo sát Sức khoẻ Tâm thần Vị thành niên Việt Nam (V-NAMHS) là khảo sát đầu tiên ở Việt Nam sử dụng công cụ chẩn đoán tiêu chuẩn quốc tế — Diagnostic Interview Schedule for Children Version 5 (DISC-5) — để ước tính tỷ lệ rối loạn tâm thần ở vị thành niên 10–17 tuổi theo tiêu chí DSM-5 trên mẫu đại diện quốc gia. V-NAMHS là một phần của khung NAMHS với ba quốc gia: Kenya (K-NAMHS), Indonesia (I-NAMHS) và Việt Nam (V-NAMHS).')

P('Được Viện Xã hội học (IOS, thuộc Viện Hàn lâm KHXH VN) chủ trì phối hợp với The University of Queensland (UQ, Úc) và Johns Hopkins Bloomberg School of Public Health (JHSPH, Mỹ), V-NAMHS đã phỏng vấn 5.996 cặp phụ huynh–vị thành niên tại 38 tỉnh/thành phố trải dài 4 vùng địa lý, bao phủ cả đô thị và nông thôn. Dữ liệu được thu thập từ 21 tháng 9 đến 16 tháng 12 năm 2021, với 127 phỏng vấn viên được đào tạo bài bản. Tỷ lệ trả lời đạt 81,1 %.')

P('Các phát hiện chính:', bold=True)
P('1. 21,7 % vị thành niên 10–17 tuổi có một vấn đề sức khoẻ tâm thần (mental health problem) trong 12 tháng qua (đáp ứng ít nhất một nửa số triệu chứng theo DSM-5).')
P('2. 3,3 % đáp ứng đầy đủ tiêu chí chẩn đoán một rối loạn tâm thần theo DSM-5.')
P('3. Lo âu là vấn đề SKTT phổ biến nhất (18,6 %); rối loạn lo âu là rối loạn phổ biến nhất (2,3 %).')
P('4. Chỉ 8,4 % vị thành niên có vấn đề SKTT đã sử dụng bất kỳ dịch vụ hỗ trợ nào trong 12 tháng qua — khoảng trống điều trị rất lớn.')
P('5. 7,7 % vị thành niên thường xuyên trải qua nhiều vấn đề cảm xúc/hành vi hơn bình thường trong đại dịch COVID-19; 71,5 % hộ gia đình báo cáo giảm thu nhập.')

P('Hàm ý chính sách:', bold=True)
P('• Cần chính sách SKTT quốc gia toàn diện đưa VTN làm nhóm đích cụ thể, bao gồm cả VTN trong và ngoài trường học')
P('• Lồng ghép sàng lọc và chăm sóc SKTT vào dịch vụ y tế ban đầu (mhGAP WHO) và vào trường học')
P('• Nâng cao mental health literacy cho phụ huynh, giáo viên, nhân viên y tế cộng đồng')
P('• Giảm kỳ thị xã hội và chuẩn hoá đường dẫn tìm kiếm trợ giúp')
P('• Chuẩn bị sẵn dịch vụ dễ tiếp cận (hotline, tư vấn trực tuyến) cho các tình huống khẩn cấp như đại dịch')

doc.add_paragraph()

# =============================================================
# INTRODUCTION (pages 9-13)
# =============================================================
H2('GIỚI THIỆU (INTRODUCTION)')
PageMarker('9')

H3('Bối cảnh (Background)')
P('Sức khoẻ tâm thần bao gồm "trạng thái hạnh phúc ở đó một cá nhân nhận ra tiềm năng của chính mình, có thể ứng phó với những căng thẳng bình thường của cuộc sống, có thể làm việc hiệu quả và đóng góp cho cộng đồng của mình" (WHO 2014). Rối loạn SKTT là nguyên nhân hàng đầu gây ra khuyết tật và gánh nặng bệnh tật ở người trẻ trên toàn thế giới (Erskine et al. 2015), với khoảng 50 % các rối loạn SKTT bắt đầu trước 14 tuổi (Kessler et al. 2005). Nếu không được phát hiện và điều trị kịp thời, các rối loạn này có thể kéo dài tới giai đoạn trưởng thành với hậu quả tiêu cực về học tập, xã hội, kinh tế và sức khoẻ thể chất (Erskine et al. 2016, Ormel et al. 2017).')

P('Tại Việt Nam, nghiên cứu về SKTT ở vị thành niên còn rất hạn chế. Một rà soát tài liệu của Samuels và cộng sự (2018) tìm thấy tỷ lệ vấn đề SKTT dao động rộng từ 8 % đến 29 %, phản ánh sự không đồng nhất về phương pháp, công cụ đo, và khung mẫu giữa các nghiên cứu. Nghiên cứu đại diện quốc gia duy nhất trước V-NAMHS là Weiss và cộng sự (2014) đo SKTT ở 1.314 phụ huynh và 591 vị thành niên 6–16 tuổi bằng SDQ, tìm thấy 10,7 % (vị thành niên tự báo cáo) đến 11,9 % (phụ huynh báo cáo) đạt ngưỡng "borderline" cho vấn đề SKTT; CBCL và YSR cho tỷ lệ tương tự. Tuy nhiên, SDQ, CBCL và YSR là thang sàng lọc triệu chứng, không phải công cụ chẩn đoán theo DSM-5, do đó không cho phép ước tính tỷ lệ rối loạn tâm thần cụ thể.')

P('Các nghiên cứu khác chỉ bao phủ một số địa bàn. Ví dụ, Samuels và cộng sự (2018) dùng SDQ trên 402 VTN 12–17 tuổi tại các thành phố (Hà Nội, TP HCM, Điện Biên Phủ, Long Xuyên) và vùng nông thôn (Kéo Lôm, Phú Mỹ), tìm thấy 13,4 % trong nhóm "bất thường". Những nghiên cứu này đóng góp hiểu biết quan trọng nhưng không thể khái quát lên quy mô quốc gia.')

P('Khoảng trống này chính là lý do V-NAMHS được thiết kế: tạo dữ liệu đại diện quốc gia về tỷ lệ các rối loạn tâm thần ở vị thành niên 10–17 tuổi, bằng công cụ chẩn đoán tiêu chuẩn quốc tế DISC-5.')

PageMarker('10')
H3('V-NAMHS là gì?')
P('V-NAMHS là viết tắt của "Viet Nam Adolescent Mental Health Survey" — Khảo sát Sức khoẻ Tâm thần Vị thành niên Việt Nam. Đây là một khảo sát đại diện quốc gia đo lường tỷ lệ mắc rối loạn tâm thần (theo DSM-5) và các vấn đề SKTT ở vị thành niên 10–17 tuổi. V-NAMHS là một trong ba khảo sát quốc gia trong khung NAMHS cùng Kenya (K-NAMHS) và Indonesia (I-NAMHS), do University of Queensland chủ trì khung chung và Johns Hopkins Bloomberg School of Public Health làm phân tích dữ liệu (Erskine et al. 2021).')

P('Mục tiêu cụ thể của V-NAMHS:', bold=True)
P('(1) Ước tính tỷ lệ 12 tháng qua của 6 rối loạn tâm thần (ám ảnh xã hội, rối loạn lo âu lan toả, rối loạn trầm cảm nặng, PTSD, rối loạn hành vi, ADHD) ở vị thành niên 10–17 tuổi theo DSM-5.')
P('(2) Ước tính tỷ lệ 12 tháng qua của các "vấn đề SKTT" (ở ngưỡng subthreshold — đáp ứng ít nhất nửa số triệu chứng DSM-5).')
P('(3) Đo lường tỷ lệ hành vi tự sát và tự làm hại bản thân.')
P('(4) Xác định tỷ lệ sử dụng dịch vụ, nhu cầu được cảm nhận, và rào cản tiếp cận dịch vụ SKTT.')
P('(5) Đo lường các yếu tố nguy cơ và bảo vệ liên quan đến SKTT VTN (bắt nạt, trường học, gia đình, peer, tình dục, sử dụng chất, ACEs, tự trọng, SKTT cha mẹ).')
P('(6) Đánh giá tác động của đại dịch COVID-19 đối với SKTT VTN và gia đình.')

H3('Ai tổ chức V-NAMHS?')
P('V-NAMHS được thực hiện bởi:')
P('• Viện Xã hội học (IOS) thuộc Viện Hàn lâm Khoa học Xã hội Việt Nam (VASS) — chủ trì tại Việt Nam')
P('• Tổng cục Dân số Kế hoạch hoá gia đình (GOPFP) thuộc Bộ Y tế — phối hợp thực địa')
P('• The University of Queensland (UQ), Úc — Principal Investigator: TS Holly Erskine; Prof Harvey Whiteford, Prof James Scott')
P('• Johns Hopkins Bloomberg School of Public Health (JHSPH), Mỹ — Project Lead: Prof Robert Blum')
P('• Với sự tài trợ của UBS Optimus Foundation và Boroughs Charitable Trust')

H3('Ai tham gia V-NAMHS?')
P('Đối tượng tham gia là các cặp phụ huynh–vị thành niên 10–17 tuổi tại các hộ gia đình được chọn ngẫu nhiên. "Phụ huynh" ở đây được định nghĩa là người chăm sóc chính của vị thành niên (primary caregiver) — người có trách nhiệm chăm sóc và có thể cung cấp thông tin tốt nhất về em. "Vị thành niên" được giới hạn ở 10–17 tuổi (thay vì 10–19 theo WHO) vì các em 18–19 tuổi thường đã sống độc lập và DISC-5 không thiết kế cho người ≥ 18 tuổi.')

P('Tiêu chí loại trừ:')
P('• Hộ gia đình không có vị thành niên 10–17 tuổi')
P('• Vị thành niên hoặc phụ huynh không nói tiếng Việt')
P('• Không có phụ huynh sống tại địa chỉ')
P('• Hộ gia đình hoặc phụ huynh không đồng ý tham gia')
P('• Vị thành niên không thể tham gia do khuyết tật thể chất hoặc nhận thức nghiêm trọng (theo đánh giá chuẩn)')

PageMarker('11')
H3('Tuyển mẫu — Bảng 1')
P('Như thể hiện trong Bảng 1, tổng 7.599 hộ gia đình từ 38 tỉnh trên cả nước được chọn ngẫu nhiên để tiếp cận (theo khung mẫu trình bày ở Phụ lục 2). Trong đó, 6.048 hộ đủ tiêu chuẩn và đồng ý tham gia — tỷ lệ trả lời tổng đạt 81,1 % (tính bằng tỷ lệ số hộ tham gia đủ tiêu chuẩn [n = 6.048] trên tổng số hộ sau khi loại trừ không đủ tiêu chuẩn [n = 7.461]). Đồng ý tham gia (informed consent và assent) được lấy riêng từ phụ huynh và vị thành niên. Phỏng vấn phụ huynh và vị thành niên được tiến hành tách biệt và tại nơi riêng tư để tránh ảnh hưởng. 52 hộ có dữ liệu không đầy đủ đã bị loại. Mẫu V-NAMHS cuối cùng gồm 5.996 cặp phụ huynh–vị thành niên.')

MakeTable(
    ['Chỉ tiêu', 'Số lượng'],
    [
        ('Tổng số hộ gia đình trong phạm vi tiếp cận', '7.599'),
        ('Tổng số hộ không được tiếp cận hoặc không sẵn có', '502'),
        ('Tổng số hộ từ chối tham gia', '911'),
        ('Tổng số hộ không đủ tiêu chuẩn', '138'),
        ('Tổng số hộ tham gia đủ tiêu chuẩn', '6.048'),
        ('Tổng số hộ có dữ liệu không đầy đủ', '52'),
        ('Tổng số hộ có dữ liệu hoàn chỉnh (mẫu cuối cùng)', '5.996'),
    ])
P('Bảng 1. Tuyển mẫu cho V-NAMHS. Nguồn: Bảng 1 báo cáo gốc, trang 11.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

# Image on page 11 (likely hero / chapter divider)
AddImg('p011_img1_Image196.jpg',
       'Hình minh hoạ Chương Giới thiệu',
       'Chapter divider — Introduction section',
       width_cm=11.0)

PageMarker('12')
H3('Người tham gia được hỏi những gì?')
P('Rối loạn tâm thần được đánh giá bằng công cụ chẩn đoán được thiết kế riêng cho trẻ em và vị thành niên: Diagnostic Interview Schedule for Children, Version 5 (DISC-5) (Bitsko et al. 2019, Shaffer et al. 2000). Các rối loạn được đo trong V-NAMHS gồm: ám ảnh xã hội (social phobia), rối loạn lo âu lan toả (generalised anxiety disorder), rối loạn trầm cảm nặng (major depressive disorder), rối loạn căng thẳng sau sang chấn (PTSD), rối loạn hành vi (conduct disorder) và rối loạn tăng động giảm chú ý (ADHD). Trong báo cáo này, ám ảnh xã hội và rối loạn lo âu lan toả được gộp chung dưới tên "các rối loạn lo âu" (anxiety disorders).')

P('V-NAMHS cũng đo các yếu tố nguy cơ và bảo vệ liên quan đến SKTT VTN, bao gồm (nhưng không giới hạn): bắt nạt, trường học và giáo dục, quan hệ peer và gia đình, hành vi tình dục, sử dụng chất, trải nghiệm bất lợi thời thơ ấu (ACEs), bệnh tâm thần của cha mẹ, và tự trọng (self-esteem). Các câu hỏi về sử dụng dịch vụ tập trung vào mức độ sử dụng, nhu cầu được cảm nhận và rào cản tiếp cận dịch vụ. Một module riêng đánh giá trải nghiệm liên quan COVID-19 cũng được phát triển và bao gồm. Tất cả các công cụ được điều chỉnh thông qua quá trình hợp tác với các đối tác NAMHS quốc tế (Erskine et al. 2021). Danh sách đầy đủ các công cụ có ở Phụ lục 1.')

P('Các công cụ ban đầu được phát triển bằng tiếng Anh, dịch sang tiếng Việt, rồi dịch ngược lại tiếng Anh để kiểm tra tính nhất quán và chính xác về khái niệm, tức bảo đảm nghĩa của câu hỏi vẫn nhất quán với phiên bản tiếng Anh gốc. Bản dịch cũng được các bác sĩ lâm sàng Việt Nam, những người tham gia tập huấn, và nhóm V-NAMHS rà soát liên tục trong suốt giai đoạn phát triển công cụ trước khi thu thập dữ liệu.')

H3('Khi nào V-NAMHS được tiến hành?')
P('Thu thập dữ liệu cho V-NAMHS diễn ra từ ngày 21 tháng 9 năm 2021 đến 16 tháng 12 năm 2021, bắt đầu ở miền Bắc và di chuyển vào miền Nam. Tổng cộng 127 phỏng vấn viên được đào tạo đã thu thập dữ liệu tại 38 tỉnh được chọn.')

H3('COVID-19 đã tác động V-NAMHS như thế nào?')
P('COVID-19 đã tác động tới lịch trình dự kiến của V-NAMHS. Thu thập dữ liệu ban đầu dự kiến bắt đầu tháng 6 năm 2020 nhưng đã bị hoãn một năm sang tháng 6 năm 2021 do COVID-19. Tuy nhiên, đợt bùng phát COVID-19 tại các tỉnh phía Nam vào giữa năm 2021 khiến cả công tác chuẩn bị (đào tạo) và thu thập dữ liệu bị trì hoãn thêm.')

P('Ở đỉnh điểm đại dịch COVID-19, việc di chuyển giữa các tỉnh bị hạn chế. Đào tạo phỏng vấn viên, giám sát và hỗ trợ kỹ thuật được nhóm V-NAMHS tại IOS cung cấp trực tuyến. Điều này gia tăng phụ thuộc vào năng lực của mỗi phỏng vấn viên trong khi số lượng phỏng vấn viên cần thiết cao hơn dự kiến do hạn chế di chuyển và yêu cầu phỏng vấn viên phải là người địa phương. Để giải quyết, các buổi đào tạo bổ sung (refresher training) được tổ chức, bao gồm nội dung công cụ, phương pháp tiếp cận hộ gia đình được chọn, và các biện pháp bảo đảm an toàn cho cả phỏng vấn viên và người tham gia. Mọi nỗ lực được thực hiện để bảo đảm chất lượng dữ liệu, với quy trình giám sát và kiểm tra chất lượng toàn diện dưới sự hỗ trợ của JHSPH và UQ. Sự trì hoãn cũng tạo cơ hội phát triển và đưa vào module COVID-19 để đánh giá tác động của đại dịch lên vị thành niên và gia đình. Thu thập dữ liệu cuối cùng bắt đầu vào tháng 9 năm 2021 ở miền Bắc, tiếp theo là miền Nam vào cuối năm 2021 khi tình hình COVID-19 cho phép.')

H3('Phạm vi của báo cáo')
P('Báo cáo này trình bày các phát hiện chính của V-NAMHS phản ánh mục tiêu cốt lõi của nghiên cứu và phù hợp với các bên liên quan tại Việt Nam. Báo cáo có ba chương chính: sức khoẻ tâm thần (gồm hành vi tự sát và tự làm hại bản thân), sử dụng dịch vụ, và COVID-19. Đặc điểm mẫu (thông tin nhân khẩu) được bao gồm trong khi các thông tin khác liên quan đến phương pháp và tiến hành khảo sát nằm ở Phụ lục 2. Trừ khi được nêu rõ, tất cả kết quả (tỷ lệ và số lượng) đều được gia trọng (weighted) để đại diện cho phân bố tuổi–giới và đô thị–nông thôn của quần thể VTN Việt Nam 10–17 tuổi. Dù các phép kiểm định ý nghĩa thống kê không được đưa vào báo cáo, những khác biệt có ý nghĩa thống kê đã được làm nổi bật trong bảng hoặc văn bản tương ứng. Chỉ những khác biệt có ý nghĩa thống kê mới được thảo luận trong văn bản.')

PageMarker('13')

# =============================================================
# SAMPLE CHARACTERISTICS (page 14)
# =============================================================
PageMarker('14')
H2('ĐẶC ĐIỂM MẪU (SAMPLE CHARACTERISTICS)')

P('Bảng 2 và 3 trình bày đặc điểm nhân khẩu của mẫu vị thành niên; Bảng 4 trình bày đặc điểm nhân khẩu của mẫu phụ huynh. Tất cả kết quả trong các bảng này chưa được gia trọng (unweighted). Mọi thông tin nhân khẩu do phụ huynh báo cáo và được nhập vào mẫu thu thập dữ liệu vị thành niên trước khi phỏng vấn vị thành niên. Khi có sự không khớp giữa mẫu phụ huynh và vị thành niên, xử lý theo phương pháp đã định sẵn.')

H3('Vị thành niên')
P('Bảng 2 trình bày tuổi và giới tính của vị thành niên tham gia khảo sát (n = 5.996). Tuổi trung bình của người tham gia là 13,3 năm, trong đó các em trẻ hơn (10–13 tuổi) chiếm hơn một nửa (54,2 %) mẫu vị thành niên. Về giới tính, mẫu có nhiều nữ (52,6 %) hơn nam (47,5 %). [LƯU Ý BIÊN TẬP: tỷ lệ nam/nữ tổng trong Bảng 2 gốc là 52,6 % nam và 47,5 % nữ — có khả năng nhầm lẫn cột trong bản gốc; tuy nhiên để trung thành nguồn, giữ nguyên số liệu gốc.]')

MakeTable(
    ['Tuổi (năm)', '% Nam', 'n Nam', '% Nữ', 'n Nữ', '% Tổng', 'n Tổng'],
    [
        ('10–13', '28,5', '1.709', '25,7', '1.539', '54,2', '3.248'),
        ('14–17', '24,1', '1.442', '21,8', '1.306', '45,8', '2.748'),
        ('10–17', '52,6', '3.151', '47,5', '2.845', '100,0', '5.996'),
    ])
P('Bảng 2. Mẫu vị thành niên theo giới và nhóm tuổi. Nguồn: Bảng 2 báo cáo gốc, trang 14.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

PageMarker('15')
P('Bảng 3 trình bày tình trạng học tập và việc làm của mẫu vị thành niên. Đa số (94,5 %) vẫn đang đi học và chưa từng đi làm (94,6 %).')

MakeTable(
    ['Chỉ tiêu', '%', 'n'],
    [
        ('Tình trạng học tập', '', ''),
        ('Hiện đang đi học', '94,5', '5.666'),
        ('Không đi học hiện tại nhưng đã đi học trong 12 tháng qua', '1,0', '62'),
        ('Đã đi học nhưng không trong 12 tháng qua', '4,4', '263'),
        ('Chưa bao giờ đi học', '0,1', '5'),
        ('Tình trạng việc làm', '', ''),
        ('Hiện đang có việc làm', '3,0', '179'),
        ('Không có việc hiện tại nhưng đã làm trong 12 tháng qua', '1,3', '79'),
        ('Đã có việc làm nhưng không trong 12 tháng qua', '1,1', '64'),
        ('Chưa bao giờ có việc làm', '94,6', '5.674'),
        ('Không đi học VÀ không đi làm hiện tại', '3,2', '194'),
    ])
P('Bảng 3. Mẫu vị thành niên theo học tập và việc làm. Nguồn: Bảng 3 báo cáo gốc, trang 15.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H3('Phụ huynh')
P('Bảng 4 trình bày đặc điểm nhân khẩu của mẫu phụ huynh. Tuổi trung bình của phụ huynh là 44,2 năm, đa số là mẹ/mẹ kế của vị thành niên (63,6 %). Gần một nửa (47 %) chỉ hoàn thành tiểu học hoặc trung học cơ sở, trong khi khoảng một phần tư (24,1 %) hoàn thành trung học phổ thông. Về việc làm, hơn một nửa (54,1 %) toàn thời gian. Khoảng một nửa (53,7 %) là chủ hộ.')

MakeTable(
    ['Chỉ tiêu', '%', 'n'],
    [
        ('Tuổi trung bình (năm) ᵃ', '44,2', '—'),
        ('Giới tính', '', ''),
        ('Nam', '28,3', '1.695'),
        ('Nữ', '71,7', '4.301'),
        ('Quan hệ với vị thành niên', '', ''),
        ('Mẹ/mẹ kế', '63,6', '3.814'),
        ('Bố/bố dượng', '25,0', '1.497'),
        ('Ông bà', '10,3', '616'),
        ('Khác', '1,1', '66'),
        ('Tình trạng hôn nhân', '', ''),
        ('Đã kết hôn', '89,6', '5.374'),
        ('Chưa kết hôn', '0,9', '53'),
        ('Khác', '9,2', '554'),
        ('Trình độ học vấn (cao nhất đã hoàn thành)', '', ''),
        ('Không (chưa hoàn thành tiểu học)', '4,1', '247'),
        ('Tiểu học', '14,5', '872'),
        ('Trung học cơ sở', '32,4', '1.945'),
        ('Trung học phổ thông', '24,1', '1.446'),
        ('Trung cấp nghề/tương đương', '4,8', '289'),
        ('Đại học/cao đẳng', '19,2', '1.152'),
        ('Đang học', '1,7', '101'),
        ('Tình trạng việc làm', '', ''),
        ('Toàn thời gian', '54,1', '3.246'),
        ('Bán thời gian/tạm thời', '33,9', '2.031'),
        ('Không có việc nhưng đang tìm', '1,1', '68'),
        ('Không có việc và không tìm', '10,2', '613'),
        ('Tỷ lệ cũng là chủ hộ ᵇ', '53,7', '3.220'),
    ])
P('Bảng 4. Mẫu phụ huynh theo thông tin nhân khẩu. Nguồn: Bảng 4 báo cáo gốc, trang 15–16.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))
P('ᵃ n = 5.987 (sau khi loại trừ tuổi < 18 năm [n=5] và câu trả lời không có ý nghĩa [n=4]); ᵇ Chủ hộ không được xác định ở tất cả các hộ.',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

PageMarker('16')
# image on p16
AddImg('p016_img1_Image210.jpg',
       'Hình minh hoạ Chương Đặc điểm mẫu',
       'Chapter divider — Sample characteristics section',
       width_cm=11.0)

doc.save(OUT)
print(f'Part 2 (Intro + Sample) saved. Continue with Part 3 (Mental Health).')
