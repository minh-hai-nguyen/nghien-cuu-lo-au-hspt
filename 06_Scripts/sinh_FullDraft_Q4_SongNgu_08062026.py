# -*- coding: utf-8 -*-
"""Full Draft bai Q4 — Latent Profile Analysis (LPA).
Song ngu Anh-Viet, ~7500 tu, tieng Viet thuan, chi tiet cao."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'FullDraft_Q4_SongNgu_v1_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5

def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def EN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def VN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def PB(en, vn):
    EN(en); VN(vn)


# ========== TITLE BLOCK ==========
TITLE('Anxiety phenotype profiles in Vietnamese lower secondary '
      'students: A latent profile analysis with risk-protective '
      'indicator integration', 14)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Các hồ sơ phân loại lo âu ở học sinh trung học cơ sở '
              'Việt Nam: Phân tích hồ sơ tiềm ẩn tích hợp các chỉ '
              'báo nguy cơ và bảo vệ')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True; r.bold = True
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Cong Thi Hang [first author]; [Nguyen Minh Duc]; '
              'Hai-Nguyen Minh\n'
              'Faculty of Psychology and Education, Hanoi National '
              'University of Education\n'
              'Target journal: Journal of Psychopathology and '
              'Behavioral Assessment OR BMC Psychiatry\n'
              'Bản thảo đầy đủ song ngữ Anh-Việt phiên bản 1 — '
              'soạn 08/06/2026 (kế hoạch nộp 2027)')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ========== ABSTRACT ==========
H1('ABSTRACT — TÓM TẮT')

H2('Background — Bối cảnh')
PB('Variable-centered analyses such as structural equation '
   'modeling (SEM, used in our preceding papers Q2 and Q3) '
   'test how risk and protective factors relate to anxiety '
   'outcomes, but cannot detect whether students cluster into '
   'distinct "phenotype profiles" — qualitatively different '
   'patterns of risk-protective configurations. Person-'
   'centered methods such as Latent Profile Analysis (LPA) '
   'complement variable-centered approaches by identifying '
   'unobserved subgroups within a population. To date, LPA-'
   'based phenotyping of adolescent anxiety has not been '
   'applied to Vietnamese lower secondary student cohorts.',
   'Các phân tích lấy biến làm trung tâm như mô hình phương '
   'trình cấu trúc (SEM, đã sử dụng trong bài Q2 và Q3 trước '
   'đó của chúng tôi) kiểm chứng cách các yếu tố nguy cơ và '
   'bảo vệ liên quan tới các kết quả lo âu, nhưng không phát '
   'hiện được liệu học sinh có gộp thành các "hồ sơ phân '
   'loại" riêng biệt hay không — các mẫu hình cấu hình nguy '
   'cơ-bảo vệ khác biệt về mặt định tính. Các phương pháp '
   'lấy con người làm trung tâm như Phân tích Hồ sơ Tiềm ẩn '
   '(LPA — Latent Profile Analysis) bổ sung cho cách tiếp '
   'cận lấy biến làm trung tâm bằng cách nhận diện các nhóm '
   'con không quan sát được trong một dân số. Cho đến nay, '
   'phân loại hồ sơ lo âu vị thành niên dựa trên LPA chưa '
   'được áp dụng cho cohort học sinh trung học cơ sở Việt Nam.')

H2('Methods — Phương pháp')
PB('Cross-sectional survey of 1,352 Vietnamese lower '
   'secondary students (same dataset as Q2 and Q3). Ten '
   'indicator variables were submitted to LPA: three risk '
   'factors (bullying victimization, academic stress, '
   'smartphone addiction), four protective factors (school '
   'membership, parental support, peer support, self-'
   'esteem), and three DSM-5 anxiety subtype scores '
   '(Generalized Anxiety, Social Anxiety, Separation '
   'Anxiety). Models with K = 2, 3, 4, 5, 6 profiles were '
   'compared using BIC (Bayesian Information Criterion), '
   'AIC (Akaike Information Criterion), entropy, and Lo-'
   'Mendell-Rubin Likelihood Ratio Test (LMR-LRT) following '
   'Nylund et al. (2007) recommendations. The optimal '
   'solution was interpreted substantively and externally '
   'validated against gender, grade, and school '
   'distributions.',
   'Khảo sát cắt ngang trên 1.352 học sinh trung học cơ sở '
   'Việt Nam (cùng bộ dữ liệu với Q2 và Q3). Mười biến chỉ '
   'báo được đưa vào LPA: ba yếu tố nguy cơ (bị bắt nạt, áp '
   'lực học tập, nghiện điện thoại), bốn yếu tố bảo vệ (gắn '
   'bó trường, hỗ trợ cha mẹ, hỗ trợ bạn bè, lòng tự trọng), '
   'và ba điểm phân loại lo âu theo DSM-5 (Lo âu Lan tỏa, '
   'Lo âu Xã hội, Lo âu Chia ly). Các mô hình với K = 2, 3, '
   '4, 5, 6 hồ sơ được so sánh sử dụng BIC (Tiêu chí Thông '
   'tin Bayes), AIC (Tiêu chí Thông tin Akaike), độ rõ ràng '
   'phân loại (entropy), và Kiểm định Tỷ số Khả năng Lo-'
   'Mendell-Rubin (LMR-LRT) theo khuyến nghị của Nylund và '
   'cs. (2007). Giải pháp tối ưu được diễn giải về mặt thực '
   'chất và kiểm chứng bên ngoài qua phân bố giới, khối lớp, '
   'và trường học.')

H2('Results — Kết quả (dự kiến — chờ phân tích)')
PB('Chúng tôi dự đoán tìm được 4-5 hồ sơ phân loại lo âu '
   'khác biệt. Dựa trên các nghiên cứu LPA về lo âu vị thành '
   'niên ở các nước khác và việc kiểm tra trực quan sơ bộ '
   'inspection of the data, we propose the following '
   'expected profiles: (i) resilient majority with low risk '
   '+ high protective + low anxiety; (ii) achievement-driven '
   'with high academic stress + low self-esteem + moderate '
   'anxiety; (iii) peer-victimization with high bullying + '
   'low school membership + high SAD; (iv) digital-isolation '
   'with high smartphone + low parental support + moderate '
   'SocAD; and possibly (v) universally elevated with high '
   'on all risk + low on all protective + severe anxiety. '
   'Profile distributions will be tested for gender, grade, '
   'and school differences.',
   'Chúng tôi dự kiến nhận diện được 4-5 hồ sơ phân loại lo '
   'âu riêng biệt. Dựa trên các nghiên cứu LPA trước đó về '
   'lo âu vị thành niên ở các quốc gia khác và việc kiểm '
   'tra trực quan sơ bộ dữ liệu của chúng tôi, chúng tôi đề '
   'xuất các hồ sơ dự kiến sau: (i) đa số kiên cường với '
   'nguy cơ thấp + bảo vệ cao + lo âu thấp; (ii) theo đuổi '
   'thành tích với áp lực học tập cao + tự trọng thấp + lo '
   'âu vừa; (iii) nạn nhân hóa bởi bạn bè với bị bắt nạt '
   'cao + gắn bó trường thấp + SAD cao; (iv) cô lập số với '
   'nghiện điện thoại cao + hỗ trợ cha mẹ thấp + SocAD vừa; '
   'và có thể (v) đồng loạt cao với nguy cơ cao tất cả + '
   'bảo vệ thấp tất cả + lo âu nặng. Phân bố hồ sơ sẽ được '
   'kiểm chứng theo giới, khối lớp, và trường.')

H2('Conclusions — Kết luận')
PB('The phenotype profiles will inform targeted prevention '
   '— different profiles likely benefit from different '
   'intervention components. Together with our preceding Q2 '
   'paper (variable-centered SEM identifying integrated '
   'mechanisms) and Q3 paper (group-centered multi-group '
   'SEM identifying gender invariance), Q4 completes a '
   '"hat-trick" three-paper research program characterizing '
   'Vietnamese adolescent anxiety from variable-, group-, '
   'and person-centered perspectives.',
   'Các hồ sơ phân loại sẽ cung cấp thông tin cho phòng '
   'ngừa có mục tiêu — các hồ sơ khác nhau có thể hưởng '
   'lợi từ các thành phần can thiệp khác nhau. Cùng với '
   'bài Q2 trước đó (SEM lấy biến làm trung tâm xác định '
   'các cơ chế tích hợp) và bài Q3 (SEM đa nhóm lấy nhóm '
   'làm trung tâm xác định bất biến theo giới), bài Q4 '
   'hoàn thành chương trình nghiên cứu ba bài "hat-trick" '
   'mô tả đặc điểm lo âu vị thành niên Việt Nam từ ba góc '
   'nhìn: lấy biến, lấy nhóm, và lấy con người làm trung '
   'tâm.')

H2('Keywords — Từ khóa')
PB('latent profile analysis; adolescent anxiety; person-'
   'centered; phenotype profiles; Vietnam; gender; risk-'
   'protective; tam giao',
   'phân tích hồ sơ tiềm ẩn (latent profile analysis); lo '
   'âu vị thành niên; lấy con người làm trung tâm; hồ sơ '
   'phân loại; Việt Nam; giới; nguy cơ-bảo vệ; tam giáo')


# ========== INTRODUCTION ==========
H1('1. INTRODUCTION — DẪN NHẬP')

H2('1.1. Variable-centered vs person-centered approaches / '
   'Cách lấy biến làm trung tâm so với cách lấy con người '
   'làm trung tâm')
PB('Most adolescent anxiety research uses variable-centered '
   'methods — correlation, regression, structural equation '
   'modeling — that estimate average relationships between '
   'variables in the population. These methods assume the '
   'sample is homogeneous: a single set of pathway '
   'coefficients applies to all individuals. Person-centered '
   'methods such as latent profile analysis (LPA), latent '
   'class analysis (LCA), and growth mixture modeling (GMM) '
   'relax this homogeneity assumption by allowing the '
   'population to contain unobserved subgroups with '
   'qualitatively different patterns. Both approaches '
   'answer complementary scientific questions: variable-'
   'centered asks "what relationships hold on average?"; '
   'person-centered asks "how do individuals cluster into '
   'distinct types?" (Bauer 2007; Lanza & Cooper 2016).',
   'Hầu hết nghiên cứu lo âu vị thành niên sử dụng các '
   'phương pháp lấy biến làm trung tâm — tương quan, hồi '
   'quy, mô hình phương trình cấu trúc — ước lượng các '
   'mối quan hệ trung bình giữa các biến trong dân số. '
   'Các phương pháp này giả định mẫu là đồng nhất: một '
   'tập hợp các hệ số đường dẫn áp dụng cho tất cả cá '
   'nhân. Các phương pháp lấy con người làm trung tâm như '
   'phân tích hồ sơ tiềm ẩn (LPA), phân tích lớp tiềm ẩn '
   '(LCA — Latent Class Analysis), và mô hình hỗn hợp '
   'tăng trưởng (GMM — Growth Mixture Modeling) nới lỏng '
   'giả định đồng nhất này bằng cách cho phép dân số '
   'chứa các nhóm con không quan sát được với các mẫu '
   'hình khác biệt về định tính. Hai cách tiếp cận trả '
   'lời các câu hỏi khoa học bổ sung lẫn nhau: lấy biến '
   'làm trung tâm hỏi "các mối quan hệ nào giữ nguyên '
   'trung bình?"; lấy con người làm trung tâm hỏi "các '
   'cá nhân gộp thành các loại khác biệt như thế nào?" '
   '(Bauer 2007; Lanza & Cooper 2016).')

H2('1.2. Why person-centered analyses matter for prevention '
   '/ Vì sao phân tích lấy con người làm trung tâm có ý '
   'nghĩa cho phòng ngừa')
PB('Anxiety disorders are common in adolescence — '
   'approximately half of all lifetime cases begin by age '
   '14 (Kessler et al. 2005), and rates are highest in '
   'girls (Hankin et al. 2007; McLean & Anderson 2009). '
   'Cross-national epidemiology shows the burden varies '
   'widely across countries (Jefferies & Ungar 2020). '
   'Public health practice often allocates intervention '
   'resources by risk profile — high-risk youth receive '
   'intensive programs, low-risk youth receive universal '
   'prevention. However, allocation by single-variable '
   'risk (e.g., "high bullying" alone) can miss important '
   'configurations: a student with high bullying but '
   'extremely high parental support may be less in need '
   'than a student with moderate bullying combined with '
   'low parental support and low self-esteem. Identifying '
   'meaningful phenotype profiles helps match the '
   'intensity and content of intervention to specific '
   'student needs, avoiding both under-treatment of '
   'high-need students and over-treatment of low-need '
   'students.',
   'Các rối loạn lo âu phổ biến ở tuổi vị thành niên — '
   'khoảng một nửa số ca trong đời khởi phát trước 14 '
   'tuổi (Kessler và cs. 2005), và tỷ lệ cao nhất ở nữ '
   '(Hankin và cs. 2007; McLean & Anderson 2009). Dịch '
   'tễ xuyên quốc gia cho thấy gánh nặng khác biệt rộng '
   'rãi giữa các quốc gia (Jefferies & Ungar 2020). '
   'Thực hành y tế công cộng thường phân bổ nguồn lực '
   'can thiệp theo hồ sơ nguy cơ — thanh thiếu niên '
   'nguy cơ cao nhận chương trình chuyên sâu, nguy cơ '
   'thấp nhận phòng ngừa phổ quát. Tuy nhiên, phân bổ '
   'theo nguy cơ đơn biến (ví dụ "bị bắt nạt cao" một '
   'mình) có thể bỏ sót các cấu hình quan trọng: một '
   'học sinh bị bắt nạt cao nhưng có sự hỗ trợ từ cha '
   'mẹ cực kỳ cao có thể ít cần can thiệp hơn một học '
   'sinh bị bắt nạt vừa kết hợp với hỗ trợ cha mẹ thấp '
   'và lòng tự trọng thấp. Nhận diện các hồ sơ phân '
   'loại có ý nghĩa giúp khớp cường độ và nội dung can '
   'thiệp với nhu cầu cụ thể của học sinh, tránh đồng '
   'thời việc thiếu can thiệp cho học sinh có nhu cầu '
   'cao và can thiệp thừa cho học sinh có nhu cầu thấp.')

H2('1.3. Research gap: no Vietnamese phenotype profiling '
   '/ Khoảng trống nghiên cứu: chưa có phân loại hồ sơ ở '
   'Việt Nam')
PB('LPA-based phenotyping of adolescent anxiety has been '
   'conducted in samples from China, the United States, '
   'Australia, and several European countries. However, '
   'no published LPA study has examined Vietnamese lower '
   'secondary students. Given the distinctive risk-'
   'protective configuration of Vietnamese adolescents — '
   'shaped by Confucian academic culture, the integrated '
   'tam giao worldview, and the rapid digital '
   'transformation of the past decade — it is plausible '
   'that Vietnamese phenotype profiles differ from those '
   'identified in Western or other East Asian samples. '
   'Filling this gap is the central aim of the present '
   'study.',
   'Phân loại hồ sơ lo âu vị thành niên dựa trên LPA đã '
   'được thực hiện trên các mẫu từ Trung Quốc, Hoa Kỳ, '
   'Úc, và một số quốc gia châu Âu. Tuy nhiên, chưa có '
   'nghiên cứu LPA công bố nào khảo sát học sinh trung '
   'học cơ sở Việt Nam. Với cấu hình nguy cơ-bảo vệ đặc '
   'trưng của vị thành niên Việt Nam — được định hình '
   'bởi văn hóa học thuật Nho giáo, thế giới quan tích '
   'hợp tam giáo, và sự chuyển đổi số nhanh trong thập '
   'kỷ qua — có thể các hồ sơ phân loại của Việt Nam '
   'khác biệt so với các hồ sơ được nhận diện ở các mẫu '
   'phương Tây hoặc Đông Á khác. Lấp đầy khoảng trống '
   'này là mục tiêu trung tâm của nghiên cứu hiện tại.')

H2('1.4. Research questions / Các câu hỏi nghiên cứu')
PB('Unlike our Q2 and Q3 papers, which test pre-registered '
   'hypotheses, Q4 asks three exploratory research '
   'questions because LPA is by design an exploratory '
   'method. Research Question 1: How many distinct anxiety '
   'phenotype profiles are present in the sample of 1,352 '
   'Vietnamese lower secondary students? Research Question '
   '2: How are the identified profiles characterized in '
   'terms of risk factor levels, protective factor levels, '
   'and anxiety subtype severity? Research Question 3: How '
   'do profile membership distributions vary by gender, '
   'grade level, and school context (urban vs suburban)?',
   'Khác với bài Q2 và Q3 — kiểm chứng các giả thuyết '
   'đăng ký trước — bài Q4 đặt ba câu hỏi nghiên cứu '
   'thám sát vì LPA về bản chất là một phương pháp thám '
   'sát. Câu hỏi nghiên cứu 1: Có bao nhiêu hồ sơ phân '
   'loại lo âu khác biệt trong mẫu 1.352 học sinh trung '
   'học cơ sở Việt Nam? Câu hỏi nghiên cứu 2: Các hồ sơ '
   'được nhận diện được đặc trưng như thế nào về mức độ '
   'các yếu tố nguy cơ, mức độ các yếu tố bảo vệ, và độ '
   'nặng của các phân loại lo âu? Câu hỏi nghiên cứu 3: '
   'Phân bố thành viên hồ sơ thay đổi như thế nào theo '
   'giới, khối lớp, và bối cảnh trường (đô thị so với '
   'ngoại ô)?')


# ========== METHODS ==========
H1('2. METHODS — PHƯƠNG PHÁP')

H2('2.1. Participants / Người tham gia')
PB('The sample comprised 1,352 Vietnamese lower secondary '
   'students (boys n = 614, 45.4%; girls n = 738, 54.6%) '
   'recruited from two schools in Hanoi: Nhat Tan (urban) '
   'and Tay Mo (suburban), during the 2024-2025 academic '
   'year. Ages 11-14, grades 6-9 (n_grade6 = 368; n_grade7 '
   '= 316; n_grade8 = 340; n_grade9 = 328). This is the '
   'same sample used in our preceding Q2 (variable-centered '
   'SEM) and Q3 (multi-group SEM) papers, enabling direct '
   'cross-paper comparison of variable-, group-, and '
   'person-centered findings.',
   'Mẫu gồm 1.352 học sinh trung học cơ sở Việt Nam (nam '
   'n = 614, 45,4%; nữ n = 738, 54,6%) tuyển từ hai trường '
   'tại Hà Nội: Nhật Tân (đô thị) và Tây Mỗ (ngoại ô), '
   'trong năm học 2024-2025. Tuổi 11-14, lớp 6-9 (n_lớp6 '
   '= 368; n_lớp7 = 316; n_lớp8 = 340; n_lớp9 = 328). Đây '
   'là cùng mẫu được sử dụng trong bài Q2 trước đó của '
   'chúng tôi (SEM lấy biến làm trung tâm) và Q3 (SEM đa '
   'nhóm), cho phép so sánh trực tiếp xuyên bài các phát '
   'hiện lấy biến, lấy nhóm, và lấy con người làm trung '
   'tâm.')

H2('2.2. Measures and indicator selection / Công cụ đo '
   'lường và lựa chọn chỉ báo')
PB('We selected 10 indicator variables for LPA based on '
   'the integrated risk-protective model from Q2 v1. Three '
   'risk factors: bullying victimization (OBVQ); academic '
   'stress (ESSA); smartphone addiction (SAS-SV). Four '
   'protective factors: school membership (PSSM); '
   'parental support (MSPSS subscale); peer support '
   '(MSPSS subscale); self-esteem (RSES). Three anxiety '
   'subtype scores from RCADS: GAD; SAD; SocAD. All 10 '
   'indicators were standardized to a T-score metric '
   '(mean = 50, SD = 10) prior to LPA to ensure '
   'comparable scaling across indicators. Full validation '
   'details for each instrument are reported in the '
   'Methods section of our Q2 v1 paper and not repeated '
   'here.',
   'Chúng tôi chọn 10 biến chỉ báo cho LPA dựa trên mô '
   'hình tích hợp nguy cơ-bảo vệ từ Q2 v1. Ba yếu tố '
   'nguy cơ: bị bắt nạt (OBVQ); áp lực học tập (ESSA); '
   'nghiện điện thoại (SAS-SV). Bốn yếu tố bảo vệ: gắn '
   'bó trường (PSSM); hỗ trợ cha mẹ (phân thang MSPSS); '
   'hỗ trợ bạn bè (phân thang MSPSS); lòng tự trọng '
   '(RSES). Ba điểm phân loại lo âu từ RCADS: GAD; SAD; '
   'SocAD. Cả 10 chỉ báo được chuẩn hóa sang thang T-'
   'score (trung bình = 50, SD = 10) trước LPA để đảm '
   'bảo các chỉ báo có cùng thang đo. Chi tiết chuẩn '
   'hóa đầy đủ cho mỗi công cụ được báo cáo trong phần '
   'Phương pháp của bài Q2 v1 của chúng tôi và không '
   'lặp lại ở đây.')

H2('2.3. Latent profile analytic strategy / Quy trình '
   'phân tích hồ sơ tiềm ẩn')

PB('The analytic strategy comprises five sequential '
   'steps. Software for primary analyses: Mplus 8.10 or '
   'higher; for sensitivity and reproducibility checks: '
   'R 4.3+ with the "tidyLPA" package version 1.1 or '
   'higher.',
   'Quy trình phân tích gồm năm bước tuần tự. Phần mềm '
   'cho phân tích chính: Mplus 8.10 trở lên; cho kiểm '
   'tra độ nhạy và khả năng tái tạo: R 4.3 trở lên với '
   'gói "tidyLPA" phiên bản 1.1 trở lên.')

PB('Step 1 — Indicator preparation. Confirm indicator '
   'distributions are approximately normal (skewness '
   'and kurtosis within ±1). Check missing data patterns '
   '— LPA estimation handles missing data via full-'
   'information maximum likelihood (FIML). Apply T-score '
   'standardization (mean = 50, SD = 10) to each '
   'indicator for cross-indicator comparability.',
   'Bước 1 — Chuẩn bị chỉ báo. Xác nhận phân bố các chỉ '
   'báo gần với chuẩn (độ lệch và độ nhọn trong khoảng '
   '±1). Kiểm tra mẫu hình dữ liệu khuyết — LPA xử lý '
   'dữ liệu khuyết qua ước lượng hợp lý cực đại thông '
   'tin đầy đủ (FIML). Áp dụng chuẩn hóa thang T-score '
   '(trung bình = 50, SD = 10) cho mỗi chỉ báo để so '
   'sánh giữa các chỉ báo.')

PB('Step 2 — Compare K-profile solutions. Estimate LPA '
   'with K = 2, 3, 4, 5, 6 profiles. Use maximum '
   'likelihood estimation with 500 random starts to '
   'avoid local maxima — the LPA likelihood surface is '
   'multimodal, and a single starting point may converge '
   'to a suboptimal solution. Each solution is estimated '
   'independently, then compared on multiple criteria.',
   'Bước 2 — So sánh các giải pháp K hồ sơ. Ước lượng '
   'LPA với K = 2, 3, 4, 5, 6 hồ sơ. Sử dụng ước lượng '
   'hợp lý cực đại với 500 điểm khởi đầu ngẫu nhiên '
   'để tránh cực đại địa phương — bề mặt hợp lý của '
   'LPA có nhiều đỉnh, và một điểm khởi đầu duy nhất '
   'có thể hội tụ về giải pháp dưới mức tối ưu. Mỗi '
   'giải pháp được ước lượng độc lập, rồi so sánh trên '
   'nhiều tiêu chí.')

PB('Step 3 — Model selection criteria following Nylund '
   'et al. (2007) recommendations. Compare four '
   'criteria: BIC (lower = better fit while penalizing '
   'complexity); AIC (lower = better fit; less '
   'penalizing than BIC); entropy (higher = clearer '
   'classification; aim for entropy ≥ 0.80 indicating '
   'most individuals are clearly assigned to one '
   'profile); Lo-Mendell-Rubin Likelihood Ratio Test '
   '(LMR-LRT; significant p-value indicates the K-'
   'profile model fits significantly better than the '
   '(K-1)-profile model). No single criterion dictates '
   'the final solution. We select K based on '
   'convergence across criteria combined with '
   'substantive interpretability of the resulting '
   'profiles.',
   'Bước 3 — Tiêu chí chọn mô hình theo khuyến nghị '
   'của Nylund và cs. (2007). So sánh bốn tiêu chí: '
   'BIC (thấp hơn = vừa hơn đồng thời phạt sự phức '
   'tạp); AIC (thấp hơn = vừa hơn; phạt ít hơn BIC); '
   'entropy (cao hơn = phân loại rõ ràng hơn; mục '
   'tiêu entropy ≥ 0,80 chỉ ra hầu hết cá nhân được '
   'gán rõ ràng vào một hồ sơ); Kiểm định Tỷ số Khả '
   'năng Lo-Mendell-Rubin (LMR-LRT; giá trị p có ý '
   'nghĩa thống kê chỉ ra mô hình K-hồ-sơ vừa hơn '
   'đáng kể so với mô hình (K-1)-hồ-sơ). Không tiêu '
   'chí nào quyết định một mình giải pháp cuối cùng. '
   'Chúng tôi chọn K dựa trên sự hội tụ qua các tiêu '
   'chí kết hợp với tính diễn giải thực chất của các '
   'hồ sơ kết quả.')

PB('Step 4 — Profile interpretation. For the selected '
   'optimal K, examine the standardized mean profile '
   'across all 10 indicators. Each profile is given a '
   'descriptive label that captures its key features '
   '(for example, "high academic stress + low self-'
   'esteem + moderate anxiety" → "achievement-driven '
   'profile"). Estimate profile prevalence (% of '
   'sample assigned to each profile).',
   'Bước 4 — Diễn giải hồ sơ. Với K tối ưu được chọn, '
   'kiểm tra trung bình chuẩn hóa của hồ sơ qua toàn '
   'bộ 10 chỉ báo. Mỗi hồ sơ được đặt một nhãn mô tả '
   'bắt được các đặc điểm chính (ví dụ "áp lực học tập '
   'cao + tự trọng thấp + lo âu vừa" → "hồ sơ theo '
   'đuổi thành tích"). Ước lượng tỷ lệ hồ sơ (% mẫu '
   'được gán vào mỗi hồ sơ).')

PB('Step 5 — External validation. Test profile '
   'membership differences against three external '
   'variables not used in profile estimation: gender '
   '(chi-square test of profile × gender '
   'distribution), grade (chi-square of profile × '
   'grade), and school (chi-square of profile × '
   'school). For each comparison, report cell '
   'percentages, standardized residuals, and Cohen\'s '
   'w effect size. Profiles with disproportionate '
   'representation of one sub-population merit '
   'special discussion.',
   'Bước 5 — Kiểm chứng bên ngoài. Kiểm chứng khác '
   'biệt thành viên hồ sơ với ba biến bên ngoài '
   'không được sử dụng trong việc ước lượng hồ sơ: '
   'giới (kiểm định chi-square về phân bố hồ sơ × '
   'giới), khối lớp (chi-square hồ sơ × khối lớp), '
   'và trường (chi-square hồ sơ × trường). Với mỗi '
   'so sánh, báo cáo phần trăm ô, phần dư chuẩn '
   'hóa, và hệ số ảnh hưởng Cohen\'s w. Các hồ sơ '
   'có đại diện không cân đối của một dân số con '
   'xứng đáng bàn luận đặc biệt.')

H2('2.4. Ethics / Đạo đức nghiên cứu')
PB('Reuse Q2 v1 ethics section. Approval [TBD — '
   'pending official letter from HNUE Ethics Committee]. '
   'Pre-registration on Open Science Framework (OSF) '
   'completed prior to analysis; for LPA, pre-'
   'registration focuses on the analytic procedure '
   '(model selection criteria thresholds) rather than '
   'specific profile predictions, since LPA is '
   'exploratory by design.',
   'Sử dụng lại phần Đạo đức của Q2 v1. Phê duyệt '
   '[CHỜ XÁC NHẬN — chờ thư chính thức từ Hội đồng '
   'Đạo đức HNUE]. Đăng ký trước trên nền tảng Open '
   'Science Framework (OSF) hoàn thành trước phân '
   'tích; với LPA, đăng ký trước tập trung vào quy '
   'trình phân tích (ngưỡng các tiêu chí chọn mô '
   'hình) thay vì các dự đoán hồ sơ cụ thể, vì LPA '
   'về bản chất là thám sát.')


# ========== RESULTS ==========
H1('3. RESULTS — KẾT QUẢ (DỰ KIẾN)')

H2('3.1. Indicator descriptives / Thống kê mô tả các '
   'chỉ báo')
PB('Table 1 presents means, SDs, skewness, kurtosis, '
   'and intercorrelations of the 10 indicator variables. '
   'We check for floor and ceiling effects, which can '
   'distort LPA estimation. Indicators with extreme '
   'skewness (|skew| > 2) or kurtosis (|kurtosis| > 7) '
   'are flagged for transformation consideration.',
   'Bảng 1 trình bày trung bình, SD, độ lệch, độ nhọn, '
   'và liên tương quan của 10 biến chỉ báo. Chúng tôi '
   'kiểm tra hiệu ứng sàn và trần, có thể làm méo việc '
   'ước lượng LPA. Các chỉ báo có độ lệch cực đoan '
   '(|độ lệch| > 2) hoặc độ nhọn cực đoan (|độ nhọn| > '
   '7) sẽ được đánh dấu để cân nhắc chuyển đổi.')

H2('3.2. LPA model selection / Chọn mô hình LPA')
PB('Table 2 reports BIC, AIC, entropy, LMR-LRT p-value, '
   'and smallest class % for K = 2, 3, 4, 5, 6 '
   'solutions. We also report convergence success rate '
   '(out of 500 random starts) for each K. Final K is '
   'determined by triangulation of criteria + '
   'interpretability. Based on prior LPA studies of '
   'adolescent mental health (which typically identify '
   '3-6 profiles), we anticipate K = 4 or K = 5 to be '
   'optimal, but the exact K will be determined by the '
   'data.',
   'Bảng 2 báo cáo BIC, AIC, entropy, giá trị p của '
   'LMR-LRT, và % lớp nhỏ nhất cho các giải pháp K = 2, '
   '3, 4, 5, 6. Chúng tôi cũng báo cáo tỷ lệ hội tụ '
   'thành công (trong 500 điểm khởi đầu ngẫu nhiên) '
   'cho mỗi K. K cuối cùng được xác định bằng '
   'triangulation các tiêu chí + tính diễn giải. Dựa '
   'trên các nghiên cứu LPA trước đó về sức khỏe tâm '
   'thần vị thành niên (thường nhận diện 3-6 hồ sơ), '
   'chúng tôi dự đoán K = 4 hoặc K = 5 sẽ tối ưu, '
   'nhưng K chính xác sẽ được dữ liệu xác định.')

H2('3.3. Profile descriptions / Mô tả các hồ sơ')
PB('Figure 1 (a line graph or radar chart) displays '
   'the standardized mean profile across 10 indicators '
   'for each identified profile. Table 3 provides '
   'descriptive labels, prevalence percentages, and '
   'indicator-specific means and SDs for each profile. '
   'For example, an expected profile may be labeled '
   '"resilient majority" with low risk indicator means '
   '(T-score < 45), high protective indicator means '
   '(T-score > 55), and low anxiety subtype means.',
   'Hình 1 (đồ thị đường hoặc biểu đồ radar) thể hiện '
   'trung bình chuẩn hóa của hồ sơ qua 10 chỉ báo cho '
   'mỗi hồ sơ được nhận diện. Bảng 3 cung cấp các nhãn '
   'mô tả, phần trăm tỷ lệ, và trung bình + SD theo '
   'chỉ báo cho mỗi hồ sơ. Ví dụ, một hồ sơ dự kiến có '
   'thể được đặt nhãn "đa số kiên cường" với trung '
   'bình các chỉ báo nguy cơ thấp (T-score < 45), '
   'trung bình các chỉ báo bảo vệ cao (T-score > 55), '
   'và trung bình các phân loại lo âu thấp.')

H2('3.4. External validation: gender, grade, school / '
   'Kiểm chứng bên ngoài: giới, khối lớp, trường')
PB('Chi-square tests examine profile membership '
   'distributions by gender, grade, and school. For '
   'each comparison, we report observed vs expected '
   'cell percentages, standardized residuals (|z| > 2 '
   'flagged as substantively different), and Cohen\'s '
   'w effect size. Any profile that disproportionately '
   'contains one sub-population (for example, '
   '"peer-victimization profile" containing 70% girls '
   'and 30% boys when sample is 54.6% girls) is '
   'discussed in detail.',
   'Các kiểm định chi-square khảo sát phân bố thành '
   'viên hồ sơ theo giới, khối lớp, và trường. Với '
   'mỗi so sánh, chúng tôi báo cáo phần trăm ô quan '
   'sát được so với kỳ vọng, phần dư chuẩn hóa (|z| '
   '> 2 đánh dấu là khác biệt thực chất), và hệ số '
   'ảnh hưởng Cohen\'s w. Bất kỳ hồ sơ nào chứa một '
   'dân số con không cân đối (ví dụ "hồ sơ nạn nhân '
   'hóa bởi bạn bè" chứa 70% nữ và 30% nam khi mẫu '
   'là 54,6% nữ) sẽ được bàn luận chi tiết.')

H2('3.5. Profile-anxiety severity comparison / So sánh '
   'độ nặng lo âu giữa các hồ sơ')
PB('ANOVA compares raw RCADS GAD, SAD, and SocAD '
   'scores across profiles, with post-hoc Tukey '
   'comparisons. Eta-squared (η²) effect sizes '
   'quantify the proportion of anxiety variance '
   'explained by profile membership. We expect '
   'profiles to differ substantially on anxiety '
   'severity (η² > 0.10) — confirming that the '
   'identified profiles have meaningful clinical '
   'relevance.',
   'ANOVA so sánh điểm thô RCADS GAD, SAD, và SocAD '
   'giữa các hồ sơ, với so sánh post-hoc Tukey. Hệ '
   'số ảnh hưởng eta-squared (η²) định lượng tỷ lệ '
   'phương sai lo âu được giải thích bởi việc thuộc '
   'hồ sơ. Chúng tôi kỳ vọng các hồ sơ khác biệt rõ '
   'rệt về độ nặng lo âu (η² > 0,10) — xác nhận các '
   'hồ sơ được nhận diện có liên quan lâm sàng có ý '
   'nghĩa.')


# ========== DISCUSSION ==========
H1('4. DISCUSSION — BÀN LUẬN')

H2('4.1. Summary of profiles / Tóm tắt các hồ sơ')
PB('We identified [TBD: K] distinct anxiety phenotype '
   'profiles in a sample of 1,352 Vietnamese lower '
   'secondary students using latent profile analysis. '
   'The profiles were [TBD: descriptive labels], with '
   'prevalences ranging from [TBD]% to [TBD]%. Profile '
   'membership was significantly associated with '
   '[TBD: gender / grade / school context]. Anxiety '
   'severity differed substantially across profiles '
   '(η² = [TBD]).',
   'Chúng tôi nhận diện được [CHỜ: K] hồ sơ phân '
   'loại lo âu riêng biệt trong mẫu 1.352 học sinh '
   'trung học cơ sở Việt Nam sử dụng phân tích hồ sơ '
   'tiềm ẩn. Các hồ sơ là [CHỜ: nhãn mô tả], với tỷ '
   'lệ dao động từ [CHỜ]% đến [CHỜ]%. Việc thuộc hồ '
   'sơ liên quan có ý nghĩa với [CHỜ: giới / khối lớp '
   '/ bối cảnh trường]. Độ nặng lo âu khác biệt rõ '
   'rệt giữa các hồ sơ (η² = [CHỜ]).')

H2('4.2. Theoretical contribution: complement to SEM '
   '(Q2/Q3) / Đóng góp lý thuyết: bổ sung cho SEM '
   '(Q2/Q3)')
PB('Our preceding Q2 paper used variable-centered SEM '
   'to identify which risk-protective pathways matter '
   'on average. Our Q3 paper used multi-group SEM to '
   'identify how those pathways differ by gender. The '
   'present Q4 paper uses LPA to identify how '
   'individual students cluster into qualitatively '
   'distinct types. The three perspectives together '
   'provide a comprehensive picture of Vietnamese '
   'adolescent anxiety that no single method achieves. '
   'Variable-centered findings tell us WHAT mechanisms '
   'matter; group-centered findings tell us FOR WHOM '
   'differently; person-centered findings tell us HOW '
   'students are typologized. As Marsh et al. (2009) '
   'argue, variable-centered and person-centered '
   'findings can converge (supporting interpretation) '
   'or diverge (suggesting heterogeneous mechanisms) '
   '— both outcomes are scientifically informative.',
   'Bài Q2 trước đó của chúng tôi sử dụng SEM lấy '
   'biến làm trung tâm để xác định các đường dẫn '
   'nguy cơ-bảo vệ nào quan trọng trung bình. Bài '
   'Q3 của chúng tôi sử dụng SEM đa nhóm để xác định '
   'các đường dẫn đó khác biệt theo giới như thế '
   'nào. Bài Q4 hiện tại sử dụng LPA để xác định các '
   'học sinh cá nhân gộp thành các loại khác biệt về '
   'định tính ra sao. Ba góc nhìn cùng nhau cung cấp '
   'bức tranh toàn diện về lo âu vị thành niên Việt '
   'Nam mà không phương pháp đơn lẻ nào đạt được. '
   'Các phát hiện lấy biến làm trung tâm cho biết '
   'NHỮNG cơ chế nào quan trọng; các phát hiện lấy '
   'nhóm làm trung tâm cho biết CHO AI khác nhau; '
   'các phát hiện lấy con người làm trung tâm cho '
   'biết các học sinh được phân loại NHƯ THẾ NÀO. '
   'Như Marsh và cs. (2009) lập luận, các phát hiện '
   'lấy biến và lấy con người làm trung tâm có thể '
   'hội tụ (hỗ trợ diễn giải) hoặc phân kỳ (gợi ý '
   'các cơ chế không đồng nhất) — cả hai kết quả '
   'đều có giá trị khoa học.')

H2('4.3. Gender-role implications: co-rumination and '
   'profile distributions / Hàm ý vai trò giới: co-'
   'rumination và phân bố hồ sơ')
PB('Some profiles may show distinctive gender '
   'distributions consistent with co-rumination '
   'patterns documented by Rose (2002) and gender-'
   'role intensification at this developmental window '
   '(Hill & Lynch 1983; Hankin et al. 2007; McLean & '
   'Anderson 2009). For example, a "peer-'
   'victimization + low-school-membership" profile '
   'may over-represent girls, paralleling our Q3 '
   'findings on the bullying → SAD pathway being '
   'stronger in girls. Conversely, an "achievement-'
   'driven" profile may over-represent students of '
   'both genders but with distinctively gendered '
   'underlying mechanisms (girls\' achievement '
   'anxiety driven by perfectionism + family '
   'expectations; boys\' achievement anxiety driven '
   'by competitive comparison).',
   'Một số hồ sơ có thể thể hiện phân bố giới đặc '
   'trưng nhất quán với các mẫu hình co-rumination '
   'được Rose (2002) ghi nhận và sự gia tăng vai '
   'trò giới ở cửa sổ phát triển này (Hill & Lynch '
   '1983; Hankin và cs. 2007; McLean & Anderson '
   '2009). Ví dụ một hồ sơ "nạn nhân hóa bởi bạn '
   'bè + gắn bó trường thấp" có thể đại diện quá '
   'mức nữ, song song với các phát hiện Q3 của '
   'chúng tôi về đường dẫn bị bắt nạt → SAD mạnh '
   'hơn ở nữ. Ngược lại, một hồ sơ "theo đuổi '
   'thành tích" có thể đại diện quá mức học sinh '
   'cả hai giới nhưng với các cơ chế ẩn dưới mang '
   'tính giới đặc thù (lo âu thành tích của nữ '
   'được dẫn dắt bởi chủ nghĩa hoàn hảo + kỳ vọng '
   'gia đình; lo âu thành tích của nam được dẫn '
   'dắt bởi so sánh cạnh tranh).')

H2('4.4. Cultural-developmental interpretation / '
   'Diễn giải văn hóa-phát triển')
PB('Reuse cultural framing from Q2 v1 Section 4.4 — '
   'tam giao (Buddhism + Confucianism + Taoism) + '
   'Confucian academic culture + co-rumination '
   'patterns. We interpret which identified profiles '
   'reflect typical Vietnamese socio-cultural '
   'niches. For example, an "achievement-driven" '
   'profile likely reflects the Confucian academic '
   'achievement script (Stankov 2010), particularly '
   'in girls who face dual pressure to excel '
   'academically AND conform to traditional female '
   'roles. A "high-peer-victimization + low-school-'
   'membership" profile reflects the breakdown of '
   'the Confucian ideal of harmonious school '
   'community — a normative ideal that, when '
   'violated, may produce particularly intense '
   'distress because it conflicts with cultural '
   'expectations.',
   'Sử dụng lại khung văn hóa từ Q2 v1 Section 4.4 '
   '— tam giáo (Phật + Nho + Đạo) + văn hóa học '
   'thuật Nho giáo + các mẫu hình co-rumination. '
   'Chúng tôi diễn giải các hồ sơ được nhận diện '
   'phản ánh các vị thế xã hội-văn hóa Việt Nam '
   'điển hình. Ví dụ một hồ sơ "theo đuổi thành '
   'tích" có thể phản ánh kịch bản thành tích học '
   'thuật Nho giáo (Stankov 2010), đặc biệt ở nữ '
   '— những người chịu áp lực kép phải xuất sắc về '
   'học tập VÀ phải tuân theo các vai trò nữ '
   'truyền thống. Một hồ sơ "bị bắt nạt từ bạn bè '
   'cao + gắn bó trường thấp" phản ánh sự đổ vỡ '
   'của lý tưởng Nho giáo về cộng đồng trường học '
   'hài hòa — một lý tưởng chuẩn mực mà khi bị vi '
   'phạm có thể tạo ra đau khổ đặc biệt mãnh liệt '
   'vì xung đột với kỳ vọng văn hóa.')

H2('4.5. Practical implications: profile-tailored '
   'intervention / Hàm ý thực tiễn: can thiệp đặc '
   'thù theo hồ sơ')
PB('Each identified profile suggests a different '
   'intervention package. For the resilient '
   'majority: universal psycho-education on stress '
   'management and emotion regulation, delivered '
   'school-wide via existing health curriculum. For '
   'the achievement-driven profile: cognitive-'
   'behavioral reframing of academic pressure + '
   'self-compassion practice + parent psycho-'
   'education on healthy expectations. For the '
   'peer-victimization profile: targeted anti-'
   'bullying program + assertiveness training + '
   'school climate intervention. For the digital-'
   'isolation profile: parental engagement program '
   '+ structured screen-time management + offline '
   'social skills training. For the severe-anxiety '
   'profile: clinical referral for individualized '
   'psychological assessment and intervention.',
   'Mỗi hồ sơ được nhận diện gợi ý một gói can '
   'thiệp khác nhau. Với đa số kiên cường: giáo '
   'dục tâm lý phổ quát về quản lý stress và điều '
   'tiết cảm xúc, triển khai toàn trường qua '
   'chương trình sức khỏe hiện có. Với hồ sơ theo '
   'đuổi thành tích: tái cấu trúc nhận thức-hành '
   'vi về áp lực học tập + thực hành lòng tự '
   'trắc ẩn + giáo dục tâm lý cho cha mẹ về kỳ '
   'vọng lành mạnh. Với hồ sơ nạn nhân hóa bởi '
   'bạn bè: chương trình chống bắt nạt có mục '
   'tiêu + huấn luyện sự quyết đoán + can thiệp '
   'môi trường trường học. Với hồ sơ cô lập số: '
   'chương trình kết nối cha mẹ + quản lý thời '
   'gian màn hình có cấu trúc + huấn luyện kỹ '
   'năng xã hội ngoại tuyến. Với hồ sơ lo âu '
   'nặng: chuyển tuyến lâm sàng để đánh giá và '
   'can thiệp tâm lý cá nhân hóa.')

H2('4.6. Limitations / Hạn chế')
PB('Four major limitations. First, cross-sectional '
   'design means profiles are snapshots; transitions '
   'between profiles over time cannot be tested. '
   'Second, LPA assumes within-profile homogeneity, '
   'but within-profile variability may still be '
   'substantial; some students assigned to a profile '
   'may not fit it well. Third, profile replicability '
   'requires testing in external samples — single-'
   'sample LPA findings are inherently provisional. '
   'Fourth, model selection criteria (LMR-LRT, BIC, '
   'AIC, entropy) may disagree, and final K '
   'selection retains an irreducible subjective '
   'element.',
   'Bốn hạn chế chính. Thứ nhất, thiết kế cắt '
   'ngang có nghĩa các hồ sơ là ảnh chụp nhanh; '
   'các chuyển tiếp giữa các hồ sơ theo thời gian '
   'không thể kiểm chứng được. Thứ hai, LPA giả '
   'định tính đồng nhất trong hồ sơ, nhưng biến '
   'thiên trong hồ sơ có thể vẫn đáng kể; một số '
   'học sinh được gán vào một hồ sơ có thể không '
   'phù hợp với nó. Thứ ba, khả năng tái tạo hồ '
   'sơ đòi hỏi kiểm chứng trên các mẫu bên ngoài '
   '— các phát hiện LPA từ một mẫu duy nhất về '
   'bản chất là tạm thời. Thứ tư, các tiêu chí '
   'chọn mô hình (LMR-LRT, BIC, AIC, entropy) có '
   'thể không nhất quán, và việc chọn K cuối '
   'cùng vẫn giữ lại một yếu tố chủ quan không '
   'thể loại bỏ.')

H2('4.7. Future directions / Hướng nghiên cứu '
   'tiếp theo')
PB('Three priorities. First, latent transition '
   'analysis (LTA) on a longitudinal extension '
   'of the cohort would test profile stability '
   'over time — do students remain in their '
   'profile, or do many transition (e.g., from '
   'resilient majority to achievement-driven '
   'during high-stakes exam years)? Second, '
   'replication of LPA findings in rural '
   'Vietnamese and ethnic-minority cohorts would '
   'test generalizability beyond urban-suburban '
   'Hanoi. Third, linking profiles to objective '
   'outcomes — school attendance, academic '
   'performance, healthcare utilization, peer '
   'sociometric ratings — would establish the '
   'external validity of the identified profiles '
   'and strengthen the case for profile-tailored '
   'intervention.',
   'Ba ưu tiên. Thứ nhất, phân tích chuyển tiếp '
   'tiềm ẩn (LTA — Latent Transition Analysis) '
   'trên một phần mở rộng dọc của cohort sẽ '
   'kiểm chứng độ ổn định hồ sơ theo thời gian '
   '— học sinh có ở lại trong hồ sơ của họ '
   'không, hay nhiều người chuyển tiếp (ví dụ '
   'từ đa số kiên cường sang theo đuổi thành '
   'tích trong các năm thi cử quan trọng)? '
   'Thứ hai, lặp lại các phát hiện LPA trên '
   'các cohort Việt Nam ở nông thôn và dân tộc '
   'thiểu số sẽ kiểm chứng tính tổng quát '
   'ngoài Hà Nội đô thị-ngoại ô. Thứ ba, liên '
   'kết các hồ sơ với các kết quả khách quan '
   '— chuyên cần, thành tích học tập, sử dụng '
   'dịch vụ y tế, xếp hạng xã hội từ bạn bè '
   '— sẽ thiết lập giá trị bên ngoài của các '
   'hồ sơ được nhận diện và củng cố lý do cho '
   'can thiệp đặc thù theo hồ sơ.')


# ========== REFERENCES ==========
H1('5. REFERENCES — TÀI LIỆU THAM KHẢO')
refs = [
    '[1] Hankin BL, Mermelstein R, Roesch L. (2007). Sex '
    'differences in adolescent depression: Stress exposure '
    'and reactivity models. Child Development 78(1):279-'
    '295. PMID: 17328705.',
    '[2] McLean CP, Anderson ER. (2009). Brave men and '
    'timid women? A review of the gender differences in '
    'fear and anxiety. Clinical Psychology Review '
    '29(6):496-505. PMID: 19541399.',
    '[3] Kessler RC, Berglund P, Demler O, Jin R, '
    'Merikangas KR, Walters EE. (2005). Lifetime '
    'prevalence and age-of-onset distributions of DSM-IV '
    'disorders in the National Comorbidity Survey '
    'Replication. Archives of General Psychiatry '
    '62(6):593-602. PMID: 15939837.',
    '[4] Rose AJ. (2002). Co-rumination in the friendships '
    'of girls and boys. Child Development 73(6):1830-'
    '1843. PMID: 12487497. DOI: 10.1111/1467-8624.00509.',
    '[5] Stankov L. (2010). Unforgiving Confucian '
    'culture: A breeding ground for high academic '
    'achievement, test anxiety and self-doubt? Learning '
    'and Individual Differences 20(6):555-563. '
    'DOI: 10.1016/j.lindif.2010.05.003.',
    '[6] Small S, Blanc J. (2021). Mental Health During '
    'COVID-19: Tam Giao and Vietnam\'s Response. Frontiers '
    'in Psychiatry 11:589618. PMID: 33536961. '
    'DOI: 10.3389/fpsyt.2020.589618.',
    '[7] Hill JP, Lynch ME. (1983). The intensification '
    'of gender-related role expectations during early '
    'adolescence. In: Brooks-Gunn J, Petersen AC (eds.), '
    'Girls at puberty. New York: Plenum Press, pp. '
    '201-228.',
    '[8] Bauer DJ. (2007). Observations on the use of '
    'growth mixture models in psychological research. '
    'Multivariate Behavioral Research 42(4):757-786. '
    'DOI: 10.1080/00273170701710338.',
    '[9] Lanza ST, Cooper BR. (2016). Latent class '
    'analysis for developmental research. Child '
    'Development Perspectives 10(1):59-64. '
    'DOI: 10.1111/cdep.12163.',
    '[10] Marsh HW, Lüdtke O, Trautwein U, Morin AJS. '
    '(2009). Classical latent profile analysis of '
    'academic self-concept dimensions: Synergy of '
    'person- and variable-centered approaches. '
    'Structural Equation Modeling 16(2):191-225. '
    'DOI: 10.1080/10705510902751010.',
    '[11] Nylund KL, Asparouhov T, Muthén BO. (2007). '
    'Deciding on the number of classes in latent class '
    'analysis and growth mixture modeling: A Monte '
    'Carlo simulation study. Structural Equation '
    'Modeling 14(4):535-569. '
    'DOI: 10.1080/10705510701575396.',
    '[12] Jefferies P, Ungar M. (2020). Social anxiety '
    'in young people: A prevalence study in seven '
    'countries. PLOS ONE 15(9):e0239133. '
    'DOI: 10.1371/journal.pone.0239133.',
    '[13] Compas BE, Jaser SS, Bettis AH, et al. '
    '(2017). Coping, emotion regulation, and '
    'psychopathology in childhood and adolescence: A '
    'meta-analysis and narrative review. Psychological '
    'Bulletin 143(9):939-991. PMID: 28616996.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Footer
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Soạn 08/06/2026 — Bản thảo đầy đủ Q4 phiên bản 1 '
              '— kế hoạch nộp 2027 (sau khi Q2 + Q3 đã công bố)')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
