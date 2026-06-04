from docx import Document
from docx.shared import Pt, Inches, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# ===== STYLE SETUP =====
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

# Set margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

def add_heading_styled(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_bold_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
    return p

def add_italic_para(text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)
    return p

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def make_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in p.runs:
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
        set_cell_shading(cell, 'D9E2F3')
    # Data rows
    for r_idx, row_data in enumerate(rows):
        for c_idx, val in enumerate(row_data):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = str(val)
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.name = 'Times New Roman'
                    run.font.size = Pt(11)
    return table

# ========== TITLE ==========
title = doc.add_heading('', level=0)
run = title.add_run('BAO CAO TONG HOP\n3 BAI NGHIEN CUU VIET NAM VE ROI LOAN LO AU')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_italic_para('Ngay bao cao: 25/03/2026. Nguon: 3 bai PDF trong thu muc papers/Viet-nam/')
doc.add_paragraph()

# ========== PHAN I ==========
add_heading_styled('I. TOM TAT TUNG BAI', 1)

# --- BAI 1 ---
add_heading_styled('Bai 1: Nguyen Thi Van (2020) - Yeu to anh huong den RLLA o HS THPT tai TP.HCM', 2)

make_table(
    ['Muc', 'Chi tiet'],
    [
        ['Tac gia', 'Nguyen Thi Van - DH KHXH&NV, DHQG TP.HCM'],
        ['Tap chi', 'Tap chi Khoa hoc Quan ly Giao duc, So 01(25), 3/2020'],
        ['ISSN', '2354-0788'],
        ['Mau', 'Sang loc 558 HS THPT; phan tich sau 90 HS co RLLA'],
        ['Cong cu', 'Thang do STAI (Spielberger), thich nghi hoa boi Nguyen Cong Khanh (2000)'],
        ['Dia diem', '4 truong THPT tai TP.HCM (noi thanh + ngoai thanh)'],
    ]
)
doc.add_paragraph()

add_bold_para('Phuong phap:')
add_para('Nghien cuu su dung thiet ke cat ngang, sang loc 558 HS THPT bang thang do lo au STAI cua Spielberger, phien ban da duoc Nguyen Cong Khanh thich nghi hoa tai Viet Nam tu nam 2000. Ket qua sang loc xac dinh 15-18.5% HS co bieu hien RLLA. Tu nhom nay, 90 HS duoc dua vao phan tich sau ve cac yeu to anh huong, su dung mau thuan tien tai 4 truong thuoc 2 khu vuc (noi thanh va ngoai thanh). Du lieu duoc phan tich bang tuong quan Pearson va hoi quy tuyen tinh da bien (Stepwise) tren SPSS.')

add_bold_para('Ket qua chinh:')
add_para('Nghien cuu xac dinh 4 nhom yeu to anh huong den RLLA: (1) hoc tap, (2) gia dinh, (3) quan he xa hoi, (4) ban than hoc sinh. Nhom yeu to "ban than hoc sinh" co tuong quan manh nhat voi thang do lo au SAS (r=0.42), giai thich 85.4% su bien thien diem lo au. Ba yeu to ban than quan trong nhat la: it giao tiep chia se (X=0.91), cam thay cuoc song that vong (X=0.90), va lo so khi gap kho khan (X=0.88). Dieu nay cho thay HS thieu ky nang ung pho la yeu to noi tai truc tiep nhat dan den RLLA.')

add_para('Nhom yeu to hoc tap chi phoi manh thu hai (r=0.37, giai thich 74%). Ap luc thi dai hoc chiem ty le TX+RTX cao nhat (56.7%), tiep theo la dinh huong nghe nghiep (51.5%) va ap luc ket qua thi (50.6%). Ap luc ky vong tu cha me ve hoc tap cung chiem 48.9%. Yeu to gia dinh noi bat: cha me so sanh con voi nguoi khac (25.8%), ky vong qua cao/thap (25.1%), ap dat la mang (21%). Yeu to xa hoi: nguoi xung quanh gay kho chiu (18.2%), thay co mang cam xuc tieu cuc vao lop (16%).')

doc.add_paragraph()

# --- BAI 2 ---
add_heading_styled('Bai 2: Tran Thi My Luong (2020) - Roi loan lo au o HS THPT Chuyen', 2)

make_table(
    ['Muc', 'Chi tiet'],
    [
        ['Tac gia', 'Tran Thi My Luong & Dang Duc Anh - Hoc vien Phu nu Viet Nam'],
        ['Tap chi', 'Tap chi Khoa hoc, Truong DH Thu Do Ha Noi, So 40/2020'],
        ['Mau', '540 HS THPT Chuyen (3 khoi 10-12), chon mau ngau nhien phan tang'],
        ['Cong cu', 'DASS-42 (sang loc) + Phieu hoi (bieu hien & nguyen nhan)'],
        ['Gioi tinh mau', 'Nam 33.5%, Nu 66.5%'],
    ]
)
doc.add_paragraph()

add_bold_para('Ty le RLLA:')
add_para('Tren 540 HS duoc sang loc, ty le mac RLLA la 14.2% (77 HS). Phan theo muc do: nhe 3.5% (19 HS), vua 7.2% (39 HS), nang 2.4% (13 HS), rat nang 1.1% (6 HS). So voi cac dieu tra dich te trong nuoc va quoc te (dao dong 8-20% o thanh thieu nien), ty le 14.2% nam o muc trung binh. Ty le RLLA muc do rat nang chi chiem 1.1%, cho thay tinh trang van trong tam kiem soat duoc.')

add_bold_para('Phat hien bat ngo - Khoi 11 cao nhat:')
add_para('Ket qua sang loc cho thay khoi 11 co ty le RLLA cao nhat (48.1% tren tong 77 HS), khong phai khoi 12 nhu gia thuyet ban dau. Khoi 10 chiem 31.1%, khoi 12 thap nhat chi 20.8%. O muc do vua va nang, so HS khoi 11 gap 2-3 lan so voi 2 khoi con lai. Ly giai: HS lop 11 chiu ap luc da chieu - ban khoan chon truong thi, nganh thi, khoi thi; quan he xa hoi phuc tap; thieu nguon ho tro thong tin huong dan nen mong lung lo lang. HS lop 12 da xac dinh duoc luc hoc va huong di, kha nang tu duy giai quyet van de tot hon.')

add_bold_para('Bieu hien the chat va tam ly:')
add_para('85.7% HS co bieu hien choang vang, chong mat, hoa mat - day la bieu hien phoi bien nhat. 76.6% kho tap trung chu y, tri nho trong rong. 62.3% roi loan giac ngu va nong ruot. Ve tam ly: 85.7% de cau gat vi chuyen khong dang va de yeu met moi. 81.8% do du khi phai dua ra quyet dinh. 80.5% lo lang ve tinh huong bi che gieu. 67.5% that vong ve ban than, cam thay vo dung bat luc. Dac biet, menh de "cam thay ban than vo dung" co ty le "Rat thuong xuyen" cao nhat - nghich ly o HS Chuyen von gioi nhung lai cam thay bat luc truoc ap luc danh tieng.')

doc.add_paragraph()

# --- BAI 3 ---
add_heading_styled('Bai 3: Tran Nguyen Ngoc (2018) - Luan an Tien si: Hieu qua dieu tri RLLALT bang lieu phap thu gian-luyen tap', 2)

make_table(
    ['Muc', 'Chi tiet'],
    [
        ['Tac gia', 'Tran Nguyen Ngoc - DH Y Ha Noi'],
        ['Huong dan', 'PGS.TS Nguyen Kim Viet'],
        ['Chuyen nganh', 'Tam than, Ma so 62720148'],
        ['Dia diem', 'Vien Suc khoe Tam than, BV Bach Mai (10/2013-10/2017)'],
        ['Mau MT1', '170 benh nhan RLLALT (F41.1 theo ICD-10), mo ta cat ngang'],
        ['Mau MT2', '99 benh nhan can thiep thu gian-luyen tap, 20 buoi/4 tuan, khong dung thuoc'],
        ['Cong cu', 'HAM-A, PSQI, Eysenck EPI, CGI'],
    ]
)
doc.add_paragraph()

add_bold_para('Dac diem lam sang (n=170):')
add_para('Benh nhan phan lon la nu (61.8%), tuoi trung binh 43.2 +/- 13.6, nhom 26-45 tuoi thuong gap nhat. Chi 22.9% den dung chuyen khoa tam than lan kham dau tien; 26.4% kham tim mach, 16.5% than kinh - phan anh nhan thuc yeu ve SKTT va ky thi benh tam than tai Viet Nam. Lo au nang theo HAM-A chiem 37.1%. Chu de lo au: gia dinh (79.4%), tai nan/benh tat (72.4%), cong viec (63.5%). Thuong gap 3 chu de lo au cung luc (40%). Trieu chung pho bien nhat: kho ngu vi lo lang (97%), bon chon (93.5%), hoi hop/tim dap nhanh (89.4%), cang thang tam than (71.7%).')

add_para('77.2% benh nhan co nhan cach huong noi-khong on dinh (Eysenck). 45.3% co sang chan tam ly, chu yeu truong dien, chu de gia dinh. 12.4% dong mac tram cam. Trieu chung khoi phat thuong gap: hoi hop/tim dap nhanh (40%), bon chon (35.3%), kho ngu (16.5%). Thoi diem lo au nang len: buoi toi (68.2%). Tan suat xuat hien trung binh 5.4 lan/tuan, thoi gian ton tai 21-31 phut moi lan.')

add_bold_para('Hieu qua dieu tri (n=99):')

make_table(
    ['Chi so', 'T0 (truoc)', 'T2 (tuan 2)', 'T4 (tuan 4)', 'p (T0-T4)'],
    [
        ['Lo au nang (HAM-A)', '45.5%', '22.2%', '11.1%', '<0.0001'],
        ['Tan suat lo au (lan/tuan)', '5.2 +/- 2.7', '3.4 +/- 2.6', '1.3 +/- 2.0', '<0.0001'],
        ['So trieu chung TB', '11.8 +/- 3.5', '9.5 +/- 3.8', '5.1 +/- 4.9', '<0.0001'],
        ['TC than kinh thuc vat', '2.5 +/- 1.0', '1.7 +/- 1.1', '0.9 +/- 1.1', '<0.0001'],
        ['Kho ngu', '96.9%', '96.9%', '75.7%', '<0.0001'],
        ['Bon chon', '96.9%', '96.9%', '52.5%', '<0.0001'],
        ['CGI - thuyen giam ro ret', '1.0%', '-', '41.4%', '<0.0001'],
    ]
)
doc.add_paragraph()

add_para('Lieu phap thu gian-luyen tap (ket hop tu am thi Schultz + tho khi cong + 6 tu the Yoga) cho thay hieu qua ro ret trong giam trieu chung lo au va cac trieu chung co the. Tuy nhien, tac dung tu tu - khong co benh nhan nao dat "cai thien ro ret" hay "cai thien rat nhieu" theo CGI. Trieu chung kho ngu va bon chon khong thay doi sau 2 tuan dau, chi cai thien tu tuan 2-4. Trieu chung tri giac sai thuc tai khong cai thien co y nghia thong ke (p=0.1574). 54.55% benh nhan tu danh gia "co hieu qua", 12.12% "khong hieu qua".')

doc.add_paragraph()

# ========== PHAN II ==========
add_heading_styled('II. DIEM NOI BAT', 1)

items = [
    ('1. Ty le RLLA o HS THPT Viet Nam kha cao (14-18.5%)',
     'Ca 2 bai ve hoc sinh deu xac nhan ty le RLLA nam trong dai 14-18.5%, phu hop voi du lieu dich te quoc te (8-20% o thanh thieu nien). Day la bang chung quan trong cho thay RLLA o HS THPT Viet Nam la van de suc khoe tam than can duoc quan tam. Cac so lieu nay duoc thu thap truoc COVID-19, nen ty le thuc te hien nay co the cao hon nhieu.'),
    ('2. Khoi 11 co RLLA cao nhat - phat hien bat ngo',
     'Bai 2 phat hien khoi 11 co ty le RLLA cao nhat (48.1%), khong phai khoi 12 nhu gia thuyet. Day la ket qua co gia tri thuc tien lon. HS lop 11 chiu ap luc da chieu: chua xac dinh duoc truong/nganh thi, quan he xa hoi phuc tap, thieu nguon ho tro. Trong khi lop 12 da on dinh hon ve dinh huong va tu duy. Ket qua nay goi y can tap trung can thiep tam ly cho HS lop 11 tu som.'),
    ('3. Nghich ly HS Chuyen - "perfectionism trap"',
     'HS Chuyen von gioi nhung lai cam thay "vo dung, bat luc" - ty le "Rat thuong xuyen" cao nhat. Ap luc danh tieng truong Chuyen tro thanh tac nhan gay lo au thay vi dong luc. Day la hien tuong "bay chu nghia hoan hao" (perfectionism trap) duoc ghi nhan ro rang trong boi canh giao duc Viet Nam. Mot so HS thuat lai: "Em la HS Chuyen ma cung khong lam duoc, em thay tet qua."'),
    ('4. Yeu to ban than > yeu to hoc tap',
     'Bai 1 chi ra nhom yeu to "ban than hoc sinh" chi phoi RLLA manh nhat (r=0.42, 85.4%), vuot xa yeu to hoc tap (r=0.37, 74%). Thieu ky nang ung pho, thieu giao tiep chia se la yeu to noi tai quan trong nhat. Dieu nay goi y rang can thiep ky nang song (kha nang giai quyet van de, giao tiep, dieu hoa cam xuc) co the hieu qua hon chi giam tai hoc tap.'),
    ('5. Luan an tien si - cong trinh bai ban nhat',
     'Bai 3 la nghien cuu dau tien tai Viet Nam danh gia hieu qua lieu phap thu gian-luyen tap tren RLLALT theo ICD-10, voi co mau 170 benh nhan. Phan tong quan benh sinh rat xuat sac, trinh bay chi tiet co che GABA, Serotonin, Norepinephrine, Neurosteroid, mach than kinh voi nhieu tham khao quoc te. Day la dong gop ly thuyet quan trong cho nghien cuu tam than hoc tai Viet Nam.'),
    ('6. Lieu phap hieu qua nhung cham',
     'Lieu phap thu gian-luyen tap giam lo au nang tu 45.5% xuong 11.1% sau 4 tuan. Tuy nhien, khong co benh nhan nao dat "cai thien ro ret" theo CGI. Phu hop voi ban chat RLLALT la benh man tinh. Bon chon va kho ngu chi cai thien tu tuan 2-4, khong thay doi 2 tuan dau. Goi y can keo dai lieu trinh >4 tuan va ket hop voi cac phuong phap khac (CBT, thuoc) cho truong hop nang.'),
    ('7. Chi 22.9% den dung chuyen khoa tam than',
     'Hau het benh nhan kham tim mach (26.4%) hoac than kinh (16.5%) truoc khi den tam than. Phan anh 3 van de: (1) nhan thuc yeu ve SKTT o ca benh nhan va nhan vien y te, (2) ky thi benh tam than, (3) trieu chung co the cua RLLALT de gay nham lan voi benh ly thuc the. Can dao tao them cho bac si da khoa ve nhan biet RLLALT.'),
]

for title_text, body_text in items:
    add_bold_para(title_text)
    add_para(body_text)

doc.add_paragraph()

# ========== PHAN III ==========
add_heading_styled('III. PHAN BIEN', 1)

add_heading_styled('Bai 1 (Nguyen Thi Van):', 2)
add_para('Mau thuan tien (n=90 cho phan phan tich yeu to) khong dai dien, co the co thien lech chon mau. Khong the tong quat hoa cho cac tinh khac ngoai TP.HCM. Thang do STAI thich nghi tu nam 2000 - da 20 nam, can cap nhat psychometric properties. Thieu phan tich da bien (multivariate) de kiem soat yeu to gay nhieu giua cac nhom. Con so 85.4% giai thich bien thien tu yeu to ban than la qua cao - can kiem tra lai phuong phap hoi quy vi co the overfitting hoac multicollinearity. Tai lieu tham khao chi co 4 nguon, qua it cho mot bai nghien cuu. Khong co nhom doi chung (HS khong co RLLA) de so sanh.')

add_heading_styled('Bai 2 (Tran Thi My Luong):', 2)
add_para('Chi nghien cuu tai 1 truong Chuyen - dac thu cao, khong dai dien cho HS THPT noi chung. Ty le nu 66.5% trong mau gay thien lech gioi, co the anh huong ket qua. Su dung DASS-42 de sang loc nhung khong neu ro diem cat (cutoff score) duoc su dung. Phan phong van sau va phieu hoi tren 77 HS: khong neu ro phuong phap phan tich thong ke. Thieu nhom doi chung (HS khong co RLLA) de so sanh bieu hien. Ket luan ve khoi 11 can duoc kiem chung tren mau lon hon, da truong. Phan ly giai nguyen nhan chua duoc ho tro bang du lieu dinh luong ma chu yeu la mo ta dinh tinh.')

add_heading_styled('Bai 3 (Tran Nguyen Ngoc - Luan an TS):', 2)
add_para('Han che lon nhat: KHONG CO NHOM CHUNG. Thiet ke truoc-sau khong co nhom doi chung khong the loai tru hieu ung placebo, hoi quy ve trung binh, hoac tac dong cua nam vien (moi truong noi tru). 41.8% benh nhan bo cuoc (170 -> 99 hoan thanh) - ty le bo cuoc rat cao, co the gay thien lech song sot (survival bias), nhung nguoi bo cuoc co the la nhom nang nhat. Khong mu (no blinding): bac si danh gia CGI cung la nguoi huong dan dieu tri, gay thien lech danh gia. Chi theo doi 4 tuan - qua ngan cho RLLALT (benh man tinh). Lieu phap chi thuc hien tai 1 co so (Vien SKTT, BV Bach Mai) nen khong tong quat hoa duoc. Tuy nhien, phan tong quan benh sinh RLLALT rat xuat sac va day la dong gop ly thuyet quan trong.')

doc.add_paragraph()

# ========== PHAN IV ==========
add_heading_styled('IV. HUONG NGHIEN CUU DE HOAN THIEN VA PHAT TRIEN', 1)

add_heading_styled('A. Mo rong nghien cuu dich te:', 2)
add_para('1. Nghien cuu cat ngang da trung tam ve RLLA o HS THPT tai Viet Nam - bao gom nong thon, thanh thi, mien nui. Mau can >3000 HS, su dung GAD-7 hoac DASS-21 (da chuan hoa quoc te). Viec su dung cong cu chuan hoa giup so sanh truc tiep voi du lieu quoc te va tang do tin cay cua ket qua.')

add_para('2. So sanh ty le RLLA truoc va sau COVID-19 o HS Viet Nam. Hien chua co du lieu nao ve van de nay. Cac nghien cuu quoc te cho thay lo au o tre em/vi thanh nien tang gap doi trong dai dich. Can co du lieu tuong tu cho Viet Nam de xay dung chinh sach.')

add_para('3. Nghien cuu doc theo doi HS tu lop 10 den lop 12 de xac nhan phat hien "khoi 11 cao nhat". Thiet ke nay cho phep theo doi su thay doi lo au theo thoi gian va xac dinh thoi diem can thiep toi uu. Can it nhat 2-3 truong (thuong + chuyen) de so sanh.')

add_heading_styled('B. Yeu to nguy co dac thu Viet Nam:', 2)
add_para('4. Vai tro cua ap luc thi cu, hoc them, ky vong gia dinh - can RCT so sanh nhom co can thiep giam ap luc vs. nhom doi chung. Viet Nam co van hoa hoc tap canh tranh cao, hoc them pho bien. Can luong hoa moi lien he giua so gio hoc them va muc do lo au. Day la de tai co tinh ung dung thuc tien cao cho chinh sach giao duc.')

add_para('5. Tac dong cua mang xa hoi len lo au o HS Viet Nam - de tai moi, chua co nghien cuu. Tren the gioi, su dung >7 nen tang MXH tang nguy co lo au gap 3 lan. Can khao sat thoi gian su dung dien thoai/MXH va tuong quan voi trieu chung lo au tren mau HS Viet Nam.')

add_para('6. Yeu to van hoa: vai tro cua "the dien" (face), chu nghia tap the, ky thi benh tam than - so sanh voi du lieu quoc te. Viet Nam co boi canh van hoa dac thu anh huong den cach bieu hien va tim kiem ho tro cho lo au. Can nghien cuu dinh tinh ket hop dinh luong.')

add_heading_styled('C. Can thiep va dieu tri:', 2)
add_para('7. RCT so sanh lieu phap thu gian-luyen tap voi CBT hoac voi thuoc (SSRI). Thieu nhom chung la han che lon nhat cua bai 3 can duoc khac phuc. Thiet ke 3 nhanh (thu gian vs CBT vs thuoc) voi follow-up 6 thang se cung cap bang chung manh nhat.')

add_para('8. Keo dai theo doi lieu phap thu gian-luyen tap den 3-6 thang. Danh gia ty le tai phat. RLLALT la benh man tinh (60% hoi phuc trong 12 nam theo Bruce et al.), nen theo doi ngan han 4 tuan chua du de ket luan ve hieu qua lau dai.')

add_para('9. Phat trien chuong trinh CBT phien ban Viet Nam cho HS THPT. Hien tai CBT chua pho bien o Viet Nam du da duoc chung minh la lieu phap hang dau tren the gioi. Can dich va thich nghi hoa protocol CBT, dao tao nhan luc, va thu nghiem tai truong hoc.')

add_para('10. Ung dung ky thuat so: App thu gian/chanh niem bang tieng Viet cho HS. Thi truong DTx toan cau dat 3.23 ty USD (2021). Can phat trien ung dung phien ban Viet voi noi dung dua tren bang chung, co danh gia hieu qua bang RCT.')

add_heading_styled('D. Phuong phap luan:', 2)
add_para('11. Su dung thiet ke RCT hoac quasi-experimental co nhom chung. Day la yeu cau toi thieu de nang cap muc do bang chung. Can tinh co mau du lon, phan bo ngau nhien, va mu danh gia (assessor blinding).')

add_para('12. Su dung cong cu do luong chuan hoa quoc te (GAD-7, PHQ-9, DASS-21) thay vi chi STAI. GAD-7 co sensitivity 0.81, specificity 0.78 qua meta-analysis 43 nghien cuu. Viec su dung cong cu chuan hoa giup so sanh truc tiep voi van ban quoc te.')

add_para('13. Can phan tich da bien (logistic regression, SEM) de xac dinh yeu to nguy co doc lap. Cac bai hien tai chu yeu dung thong ke mo ta va tuong quan don bien, chua kiem soat yeu to gay nhieu. SEM se cho phep kiem dinh mo hinh nguyen nhan-ket qua phuc tap hon.')

doc.add_paragraph()

# ========== PHAN V ==========
add_heading_styled('V. DANH SACH BAI CAN DOC LIEN QUAN', 1)

add_heading_styled('A. Tong quan & ly thuyet nen tang:', 2)
make_table(
    ['#', 'Tac gia', 'Nam', 'Tap chi', 'Noi dung'],
    [
        ['1', 'Penninx et al.', '2021', 'The Lancet (IF~169)', 'Tong quan RLLA hang dau, moi nhat'],
        ['2', 'Szuhany & Simon', '2022', 'JAMA (IF~157)', 'Tong quan RLLA toan dien'],
        ['3', 'Craske et al.', '2017', 'Nature Rev Dis Primers', 'Primer RLLA nen tang (~1500 citations)'],
    ]
)
doc.add_paragraph()

add_heading_styled('B. Lo au o hoc sinh (lien quan bai 1 & 2):', 2)
make_table(
    ['#', 'Tac gia', 'Nam', 'Tap chi', 'Noi dung'],
    [
        ['4', 'Racine et al.', '2021', 'JAMA Pediatrics', 'Lo au 20.5%, tram cam 25.2% o TE trong COVID (~3000 cit)'],
        ['5', 'Samji et al.', '2023', 'JAMA Pediatrics', 'Thay doi lo au TE truoc vs trong COVID'],
        ['6', 'Xu Q et al.', '2021', 'J Affect Disord', 'Lo au 373,216 HS Trung Quoc trong COVID'],
        ['7', 'Chen Z et al.', '2023', 'BMC Psychiatry', 'Lo au HS THCS Trung Quoc - PHQ-9, GAD-7'],
        ['8', 'Qiu Z et al.', '2022', 'Front Psychology', 'Phong cach nuoi day & lo au HS THCS TQ'],
    ]
)
doc.add_paragraph()

add_heading_styled('C. Lieu phap thu gian & mindfulness (lien quan bai 3):', 2)
make_table(
    ['#', 'Tac gia', 'Nam', 'Tap chi', 'Noi dung'],
    [
        ['9', 'Papola et al.', '2024', 'JAMA Psychiatry', 'Network meta-analysis CBT cho GAD'],
        ['10', 'Nature Mental Health', '2023', 'Nature Mental Health', 'MBI hieu qua tuong duong CBT'],
        ['11', 'Manzoni et al.', '2008', 'BMC Psychiatry', 'Meta-analysis thu gian cho lo au'],
        ['12', 'Singh et al.', '2023', 'Br J Sports Med', 'Van dong the chat cho lo au/tram cam'],
    ]
)
doc.add_paragraph()

add_heading_styled('D. Cong cu sang loc & boi canh VN:', 2)
make_table(
    ['#', 'Tac gia', 'Nam', 'Tap chi', 'Noi dung'],
    [
        ['13', 'GAD-7 Meta-analysis', '2025', 'J Affect Disord', 'Sensitivity 0.81, specificity 0.78'],
        ['14', 'AACAP Guideline', '2020', 'JAACAP', 'Huong dan dieu tri lo au tre em'],
        ['15', 'Ho Huu Tinh et al.', '2010', 'TC Y hoc TP HCM', 'Stress lo au HS cap 3 Viet Nam'],
        ['16', 'Nguyen Phuoc Binh', '2010', 'Luan van CKII, DH Y HN', 'Dac diem LS cua RLLALT'],
    ]
)

doc.add_paragraph()
add_italic_para('Bao cao nay duoc viet ngay 25/03/2026. Dua tren phan tich 3 bai PDF trong thu muc papers/Viet-nam/.')

# Save
output_path = r'c:\Users\HLC\OneDrive\read_books\Lo-au\BAO_CAO_3_BAI_VIET_NAM.docx'
doc.save(output_path)
print(f'Saved to: {output_path}')
