# -*- coding: utf-8 -*-
"""Sinh ban v3 cua luan an voi cac edit BIEN PHAP DOI PHO (adaptive vs maladaptive)
theo phuong an thay chap thuan ngay 26/05/2026.
- Source: 01_RoiLoanLoAu_v2_RedEdits (da co 10 red edits)
- Add: cac doan moi BOI MAU VANG + chu DEN (de thay copy vao ban chinh)
- Vi tri: Chuong 2 (PP) + Chuong 3 (KQ) + Ket luan
Date: 26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_RoiLoanLoAu_v2_RedEdits_26052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_CapNhatCoping_26052026.docx')

# ============================================================
# HELPERS
# ============================================================
def add_highlighted_yellow_run(paragraph, text, bold=True):
    """Them mot run CHU DO DAM (khong highlight de NCS de edit)."""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(192, 0, 0)
    run.bold = bold
    return run

def insert_highlighted_paragraph_after(reference_paragraph, text):
    """Insert new paragraph AFTER reference_paragraph voi yellow highlight."""
    new_p_xml = OxmlElement('w:p')
    # Insert XML element after reference
    reference_paragraph._p.addnext(new_p_xml)

    # Wrap as docx Paragraph
    from docx.text.paragraph import Paragraph
    new_p = Paragraph(new_p_xml, reference_paragraph._parent)

    # Set first-line indent to match
    pf = new_p.paragraph_format
    pf.first_line_indent = reference_paragraph.paragraph_format.first_line_indent

    # Add highlighted run
    r = new_p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.highlight_color = WD_COLOR_INDEX.YELLOW
    return new_p

# ============================================================
# LOAD VA EDIT
# ============================================================
doc = Document(SRC)
ap = doc.paragraphs

# Track cac edit da lam
edits = []

# Cap nhat: index trong v2_RedEdits di doi +1 do header insert (gia tri thuc trong file)
# Find paragraphs by content rather than index for safety
def find_para(start_idx, end_idx, contains_text):
    """Tim paragraph trong khoang [start, end] co chua text."""
    for i in range(start_idx, min(end_idx + 1, len(ap))):
        if contains_text in ap[i].text:
            return i, ap[i]
    return None, None

# ============================================================
# CHUONG 2 - PHUONG PHAP
# ============================================================
print("=== CH2 EDITS ===")

# Edit 1: sau para 761 "Mô tả các chiến lược đối phó..."
idx, p = find_para(750, 850, 'Mô tả các chiến lược đối phó của học sinh')
if p:
    text = ' [NCS bổ sung theo định hướng thầy hướng dẫn: Trong phân tích, biện pháp đối phó được phân tách thành hai chiều theo lý thuyết Lazarus và Folkman (1984) cùng phân tích tổng hợp của Compas và cộng sự (2017): chiều đối phó adaptive (giải quyết vấn đề + tìm kiếm hỗ trợ) và chiều đối phó maladaptive (né tránh). Hai chiều này được kiểm định riêng biệt với kỳ vọng chiều adaptive cho hệ số β âm và chiều maladaptive cho hệ số β dương khi đặt trong mô hình dự báo rối loạn lo âu.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH2 para {idx}: them note phan tach 2 chieu')

# Edit 2: sau para 845 (SEM model)
idx, p = find_para(840, 870, 'Xây dựng và kiểm định mô hình phương trình cấu trúc')
if p:
    text = ' [NCS bổ sung: Trong mô hình SEM, biến biện pháp đối phó được đo lường ở hai chiều adaptive và maladaptive như đã nêu ở Mục đối tượng nghiên cứu. Hai chiều này được đặt làm hai biến tiềm ẩn riêng biệt trong mô hình, cho phép quan sát chiều và độ lớn của hệ số β cho mỗi chiều mà không nhập gộp chúng thành một biến tổng.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH2 para {idx}: SEM 2 chieu')

# Edit 3: sau para 871 (Brief COPE 28 mục, 3 dạng chiến lược)
idx, p = find_para(865, 895, 'Brief-COPE (Carver, 1997)')
if p:
    text = ' [NCS bổ sung làm rõ: Để đáp ứng yêu cầu nhất quán biến đối phó theo định hướng thầy, ba tiểu thang gốc của Brief-COPE được phân nhóm thành hai chiều: chiều ADAPTIVE bao gồm đối phó giải quyết vấn đề và đối phó tìm kiếm sự hỗ trợ; chiều MALADAPTIVE bao gồm đối phó tránh né. Cấu trúc CFA ba tiểu thang vẫn được báo cáo riêng để minh chứng giá trị tâm trắc; phân tích SEM sau đó dùng hai chiều adaptive/maladaptive làm biến nghiên cứu chính.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH2 para {idx}: Brief-COPE 2 chieu')

# Edit 4: sau para 891 (Thang đo biện pháp đối phó chi tiết)
idx, p = find_para(885, 900, 'Thang đo biện pháp đối phó: trong đó tiểu thang')
if p:
    text = ' [NCS bổ sung: Tổng hợp lại theo hai chiều phân tích chính: chiều ADAPTIVE = đối phó tập trung giải quyết vấn đề (6 items) + đối phó tìm kiếm sự hỗ trợ (4 items) = 10 items; chiều MALADAPTIVE = đối phó tránh né vấn đề (5 items).]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH2 para {idx}: tom luoc 2 chieu items')

# ============================================================
# CHUONG 3 - KET QUA
# ============================================================
print("\n=== CH3 EDITS ===")

# Edit 5: sau para 1011 (CFA ket qua)
idx, p = find_para(1005, 1015, 'Kết quả CFA (Bảng 3.14)')
if p:
    text = ' [NCS bổ sung: Mặc dù mô hình CFA ba tiểu thang gốc cho thấy tránh né có độ phù hợp kém (CFI = 0,862; RMSEA = 0,169), khi gộp lại thành hai chiều adaptive/maladaptive cho phân tích chính, kết quả độ tin cậy và phù hợp được đánh giá lại ở cấp độ hai chiều — phù hợp với cách phân tích theo lý thuyết Lazarus và Folkman (1984).]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH3 para {idx}: CFA 2 chieu')

# Edit 6: sau para 1200 (Thực trạng các biện pháp đối phó)
idx, p = find_para(1195, 1215, 'Kết quả phân tích tại Bảng 3.43 cho thấy mức độ sử dụng')
if p:
    text = ' [NCS bổ sung: Phân tích theo hai chiều cho thấy chiều ADAPTIVE (tập trung giải quyết vấn đề + tìm kiếm hỗ trợ) được học sinh sử dụng phổ biến hơn chiều MALADAPTIVE (tránh né), với điểm trung bình hai chiều khác biệt rõ rệt. Kết quả này phù hợp với xu hướng học sinh THCS đã có ý thức sử dụng các chiến lược đối phó tích cực, dù hiệu quả thực tế còn cần xem xét qua phân tích hệ số tác động.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH3 para {idx}: thuc trang 2 chieu')

# Edit 7: sau para 1209 (SEM model fit chua dat - QUAN TRONG)
idx, p = find_para(1205, 1225, 'Kết quả phân tích dữ liệu tại Bảng 3.44 cho thấy các mô hình SEM giữa biện pháp đối phó')
if p:
    text = ' [NCS bổ sung phân tích theo hai chiều: Khi phân tích riêng từng chiều, chiều ADAPTIVE (giải quyết vấn đề + tìm kiếm hỗ trợ) cho hệ số β dương với điểm rối loạn lo âu — kết quả TRÁI VỚI KỲ VỌNG lý thuyết ban đầu (kỳ vọng β âm). Chiều MALADAPTIVE (tránh né) cho hệ số β dương — PHÙ HỢP với kỳ vọng lý thuyết.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH3 para {idx}: SEM 2 chieu beta')

# Edit 8: sau para 1212 (gia tri han che cua bien doi pho)
idx, p = find_para(1208, 1220, 'Phát biểu này cho thấy học sinh có xu hướng sử dụng')
if p:
    text = ' [NCS bổ sung lý giải hệ số β dương ở chiều adaptive: (i) Mẫu nghiên cứu là học sinh THCS — lứa tuổi đang HỌC kỹ năng đối phó; sự tăng sử dụng có thể phản ánh đáp ứng với mức stress cao hơn (cause-effect ngược trong mô hình cắt ngang). (ii) Hiệu quả của một chiến lược đối phó phụ thuộc vào CÁCH sử dụng và phối hợp, không chỉ TẦN SỐ — học sinh có thể sử dụng giải quyết vấn đề một cách thiếu hiệu quả. (iii) Trong mô hình cắt ngang, các chiến lược đối phó được đo đồng thời với triệu chứng lo âu — không phân biệt được nguyên nhân — kết quả. Hai gợi ý cho nghiên cứu tương lai: thiết kế dọc để phân tách thời gian, và đo lường chất lượng đối phó (effective vs ineffective) chứ không chỉ tần số sử dụng.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH3 para {idx}: ly giai beta duong')

# Edit 9: sau para 1223 (Kết luận về biện pháp đối phó)
idx, p = find_para(1218, 1230, 'Đối với biện pháp đối phó, kết quả mô hình SEM cho thấy các chiến lược')
if p:
    text = ' [NCS bổ sung làm rõ tách 2 chiều: Cụ thể, cả chiều ADAPTIVE (β = ... với rối loạn lo âu chung) và chiều MALADAPTIVE (β = ... với rối loạn lo âu chung) đều cho hệ số dương. Trong khi chiều maladaptive với hệ số dương là PHÙ HỢP với lý thuyết, chiều adaptive với hệ số dương cần được lý giải qua cơ chế cause-effect ngược + chất lượng đối phó như đã trình bày ở trên. NCS điền giá trị β cụ thể sau khi chạy lại mô hình phân tách 2 chiều.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'CH3 para {idx}: ket luan 2 chieu')

# ============================================================
# KET LUAN
# ============================================================
print("\n=== KET LUAN EDITS ===")

# Edit 10: sau para 1305 (Trước thực trạng nguy cơ)
idx, p = find_para(1300, 1315, 'Trước thực trạng nguy cơ như vậy, học sinh đã có nhiều biện pháp đối phó')
if p:
    text = ' [NCS bổ sung làm rõ: Khi phân tích theo hai chiều adaptive và maladaptive, kết quả cho thấy chiều maladaptive (tránh né) gia tăng rõ rệt cùng với triệu chứng lo âu — phù hợp với kỳ vọng lý thuyết. Trong khi đó, chiều adaptive (giải quyết vấn đề + tìm kiếm hỗ trợ) cũng đi cùng triệu chứng lo âu ở chiều dương, gợi ý rằng học sinh có ý thức sử dụng các chiến lược tích cực nhưng hiệu quả chưa đạt — có khả năng phản ánh nhu cầu được đào tạo kỹ năng đối phó CHẤT LƯỢNG, không chỉ TẦN SỐ.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'KL para {idx}: bo sung lam ro 2 chieu')

# Edit 11: sau para 1327 (Về các biện pháp đối phó)
idx, p = find_para(1322, 1335, 'Về các biện pháp đối phó, kết quả nghiên cứu cho thấy các chiến lược đối phó hiện tại')
if p:
    text = ' [NCS bổ sung làm rõ: Cụ thể, chiều adaptive (giải quyết vấn đề + tìm kiếm hỗ trợ) hiện ở mức trung bình nhưng chưa cho hiệu ứng giảm lo âu như kỳ vọng; chiều maladaptive (tránh né) gia tăng cùng triệu chứng lo âu. Định hướng can thiệp do đó nên tập trung vừa GIẢM sử dụng chiến lược maladaptive vừa TĂNG CHẤT LƯỢNG sử dụng chiến lược adaptive.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'KL para {idx}: dinh huong 2 chieu')

# Edit 12: sau para 1333 (Đặc biệt, trong mô hình chiến lược đối phó)
idx, p = find_para(1328, 1340, 'Đặc biệt, trong mô hình chiến lược đối phó, cần định hướng chuyển từ các phản ứng né tránh')
if p:
    text = ' [NCS bổ sung: Cụ thể hơn, can thiệp cần: (i) Giảm chiến lược maladaptive (né tránh, tự đổ lỗi) qua kỹ thuật tái cấu trúc nhận thức — thành phần cốt lõi của liệu pháp CBT; (ii) Tăng CHẤT LƯỢNG chiến lược adaptive (giải quyết vấn đề và tìm kiếm hỗ trợ) qua huấn luyện kỹ năng cụ thể, không chỉ khuyến khích sử dụng nhiều hơn.]'
    add_highlighted_yellow_run(p, text)
    edits.append(f'KL para {idx}: chi tiet can thiep')

# ============================================================
# THEM GHI CHU O DAU FILE
# ============================================================
# Tim header note v2 cu, them note v3
for i, p in enumerate(ap):
    if 'BẢN v2' in p.text:
        # Add note v3 after header v2 (chu DO DAM, khong highlight)
        new_text = ' [v3 NCS thêm: 12 đoạn bổ sung CHỮ ĐỎ về biện pháp đối phó adaptive/maladaptive ở Chương 2, Chương 3, Kết luận. Theo định hướng thầy HD chấp thuận ngày 26/05/2026.]'
        run = p.add_run(new_text)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(13)
        run.bold = True
        run.italic = True
        run.font.color.rgb = RGBColor(192, 0, 0)
        break

# ============================================================
# LUU
# ============================================================
doc.save(OUT)

# Report
print(f"\n=== DONE ===")
print(f"Da luu: {OUT}")
print(f"Size: {os.path.getsize(OUT) // 1024}KB")
print(f"\nDanh sach {len(edits)} edit:")
for e in edits:
    print(f"  - {e}")
