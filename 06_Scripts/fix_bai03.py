# -*- coding: utf-8 -*-
"""
Sửa bài-03 "Môi trường gia đình và rối loạn lo âu ở HSTHCS" (Công Thị Hằng)
theo 6 ý phản hồi vòng 1 của Tạp chí Khoa học Giáo dục Việt Nam (VJES).

6 ý phản hồi:
1. Rà soát đúng quy định mỗi mục + sửa hết viết tắt (trừ thuật ngữ quốc tế)
2. Gửi mẫu phiếu khảo sát (việc tác giả — không sửa trong bài)
3. Việt hóa trích dẫn: "et al." → "và cộng sự"; "&" → "và"
4. Cấu trúc lại 5 mục: 1.Đặt vấn đề 2.PPNC 3.Kết quả 4.Thảo luận 5.Kết luận
5. Thêm tuyên bố trách nhiệm + xung đột lợi ích
6. Khai báo AI trong Phương pháp nghiên cứu
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import copy
import re

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/bài-03"
IN_NAME = [f for f in os.listdir(ROOT) if f.endswith('.docx') and not f.startswith('~')][0]
IN_PATH = os.path.join(ROOT, IN_NAME)
OUT_PATH = os.path.join(ROOT, "Công Thị Hằng_v2_đã sửa.docx")


def set_run_font(run, font="Times New Roman"):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def replace_para_text(para, new_text):
    """Replace all run text in paragraph with new_text (first run keeps formatting)."""
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def vietnamize_citations(text):
    """Việt hóa trích dẫn trong text: et al. → và cộng sự; & → và."""
    # "et al.," / "et al." → "và cộng sự,"
    text = re.sub(r"\bet\s+al\.\s*,", "và cộng sự,", text)
    text = re.sub(r"\bet\s+al\.", "và cộng sự", text)
    # "A, B, & C" → "A, B và C"  (dấu & trong trích dẫn)
    # Replace ", &" → " và" and " & " → " và "
    text = text.replace(", &", " và")
    text = text.replace(" & ", " và ")
    return text


# ============================================================
# 1. CITATION VIETNAMIZATION + ABBREVIATION FIX (toàn văn)
# ============================================================
ABBREV_FIXES = [
    # Bỏ viết tắt GAD — giữ cụm tiếng Việt đầy đủ
    ("Rối loạn lo âu (Generalized Anxiety Disorder - GAD) được định nghĩa",
     "Rối loạn lo âu tổng quát (generalized anxiety disorder) được định nghĩa"),
    # APA → tên đầy đủ lần đầu xuất hiện trong thân bài
    # (APA, 2013) giữ — APA là tổ chức, theo APA 7 in-text dùng tên tổ chức.
    # Lần đầu: "Hiệp hội Tâm thần học Hoa Kỳ (American Psychiatric Association, 2013)"
    # THCS trong tóm tắt VN
    ("cảm nhận quan hệ tâm lý của học sinh THCS.",
     "cảm nhận quan hệ tâm lý của học sinh trung học cơ sở."),
]

# Sửa dấu thập phân kiểu Anh trong ghi chú bảng → kiểu Việt
DECIMAL_FIXES = [
    ("** p <0.001)", "** p < 0,001)"),
    ("** p <0.01)", "** p < 0,01)"),
    ("p <0.001", "p < 0,001"),
    ("p <0.01", "p < 0,01"),
    ("p < 0.05", "p < 0,05"),
]


def fix_all_paragraph_text(doc):
    """Áp dụng việt hóa trích dẫn + sửa viết tắt + sửa thập phân cho mọi đoạn."""
    n_cite = 0
    for para in doc.paragraphs:
        original = para.text
        if not original.strip():
            continue
        new = original
        # Việt hóa trích dẫn — chỉ áp dụng cho thân bài, KHÔNG cho danh mục TLTK
        # (TLTK giữ nguyên định dạng tên tác giả gốc)
        new = vietnamize_citations(new)
        # Sửa viết tắt cụ thể
        for old, rep in ABBREV_FIXES:
            if old in new:
                new = new.replace(old, rep)
        # Sửa thập phân
        for old, rep in DECIMAL_FIXES:
            if old in new:
                new = new.replace(old, rep)
        if new != original:
            replace_para_text(para, new)
            n_cite += 1
    return n_cite


def fix_table_notes(doc):
    """Sửa dấu thập phân trong các ô bảng (ghi chú)."""
    n = 0
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    orig = para.text
                    if not orig.strip():
                        continue
                    new = orig
                    for old, rep in DECIMAL_FIXES:
                        new = new.replace(old, rep)
                    if new != orig:
                        replace_para_text(para, new)
                        n += 1
    return n


def find_para(doc, text_start):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip().startswith(text_start):
            return i
    return -1


def find_para_contains(doc, substr):
    for i, p in enumerate(doc.paragraphs):
        if substr in p.text:
            return i
    return -1


def move_element_before(src_para, target_para):
    """Di chuyển src_para element ra trước target_para element."""
    target_para._element.addprevious(src_para._element)


def insert_heading_before(doc, target_idx, text, style_ref_idx):
    """Chèn 1 đoạn heading mới trước target_idx, copy format từ style_ref_idx."""
    ref_p = doc.paragraphs[style_ref_idx]
    target_p = doc.paragraphs[target_idx]
    new_el = copy.deepcopy(ref_p._element)
    for child in list(new_el):
        new_el.remove(child)
    target_p._element.addprevious(new_el)
    new_para = doc.paragraphs[target_idx]
    r = new_para.add_run(text)
    r.bold = True
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    set_run_font(r)
    return new_para


def insert_para_before(doc, target_idx, text, size=12, indent=1.0, bold=False, italic=False):
    target_p = doc.paragraphs[target_idx]
    new_el = copy.deepcopy(target_p._element)
    for child in list(new_el):
        new_el.remove(child)
    target_p._element.addprevious(new_el)
    new_para = doc.paragraphs[target_idx]
    pf = new_para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    if indent:
        pf.first_line_indent = Cm(indent)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = new_para.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    set_run_font(r)
    return new_para


def main():
    print(f"Input: {IN_PATH}")
    doc = Document(IN_PATH)

    # ---- BƯỚC 1: Việt hóa trích dẫn + sửa viết tắt + thập phân (toàn văn) ----
    # LƯU Ý: phải làm TRƯỚC khi đụng cấu trúc để index ổn định
    n = fix_all_paragraph_text(doc)
    print(f"[1] Việt hóa trích dẫn + sửa viết tắt/thập phân: {n} đoạn")
    nt = fix_table_notes(doc)
    print(f"    Sửa thập phân trong bảng: {nt} ô")

    # ---- BƯỚC 2: Sửa tiêu đề về dạng chữ thường (không in hoa toàn bộ) ----
    idx_t1 = find_para(doc, "MÔI TRƯỜNG GIA ĐÌNH")
    idx_t2 = find_para(doc, "RỐI LOẠN LO ÂU Ở HỌC SINH")
    if idx_t1 >= 0:
        replace_para_text(doc.paragraphs[idx_t1], "Môi trường gia đình và rối loạn lo âu")
    if idx_t2 >= 0:
        replace_para_text(doc.paragraphs[idx_t2], "ở học sinh trung học cơ sở")
    print("[2] Tiêu đề VN → chữ thường")

    # ---- BƯỚC 3: Bổ sung từ khóa (hiện 3 → thêm thành 5) ----
    idx_kw_vn = find_para(doc, "Từ khóa:")
    if idx_kw_vn >= 0:
        replace_para_text(doc.paragraphs[idx_kw_vn],
                          "Từ khóa: môi trường gia đình; rối loạn lo âu tổng quát; "
                          "học sinh trung học cơ sở; hoạt động gia đình; quan hệ gia đình.")
    idx_kw_en = find_para(doc, "Keywords:")
    if idx_kw_en >= 0:
        replace_para_text(doc.paragraphs[idx_kw_en],
                          "Keywords: family environment; generalized anxiety disorder; "
                          "middle school students; family activities; familial relationships.")
    print("[3] Từ khóa VN + EN: 3 → 5")

    # ---- BƯỚC 4: Cấu trúc lại 5 mục ----
    # 4a. "1. Mở đầu" → "1. Đặt vấn đề"
    idx_modau = find_para(doc, "1. Mở đầu")
    if idx_modau >= 0:
        replace_para_text(doc.paragraphs[idx_modau], "1. Đặt vấn đề")
        print("[4a] '1. Mở đầu' → '1. Đặt vấn đề'")

    # 4b. Di chuyển 3 đoạn Cơ sở lý thuyết lên trước đoạn 'khoảng trống' của §1
    # Đoạn theory bắt đầu bằng "Rối loạn lo âu tổng quát (generalized" (đã sửa ở bước 1)
    idx_theory1 = find_para_contains(doc, "Rối loạn lo âu tổng quát (generalized anxiety disorder) được định nghĩa")
    idx_gap = find_para_contains(doc, "Mặc dù đã có các nghiên cứu đề cập đến lo âu ở trẻ em")
    if idx_theory1 >= 0 and idx_gap >= 0 and idx_theory1 > idx_gap:
        # 3 đoạn theory liên tiếp: idx_theory1, +1, +2
        theory_paras = [doc.paragraphs[idx_theory1],
                        doc.paragraphs[idx_theory1 + 1],
                        doc.paragraphs[idx_theory1 + 2]]
        gap_para = doc.paragraphs[idx_gap]
        for tp in theory_paras:
            move_element_before(tp, gap_para)
        print("[4b] Di chuyển 3 đoạn Cơ sở lý thuyết lên trước đoạn khoảng trống")

    # 4c. Xóa heading "2. Nội dung" và "2.1. Cơ sở lý thuyết"
    for h in ["2. Nội dung", "2.1. Cơ sở lý thuyết"]:
        idx = find_para(doc, h)
        if idx >= 0:
            p = doc.paragraphs[idx]
            p._element.getparent().remove(p._element)
            print(f"[4c] Xóa heading '{h}'")

    # 4d. Chèn heading "2. Phương pháp nghiên cứu" trước "Khách thể nghiên cứu"
    idx_kt = find_para_contains(doc, "Khách thể nghiên cứu")
    # Tìm chính xác heading "2.2. Khách thể nghiên cứu"
    idx_kt = -1
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() in ("2.2. Khách thể nghiên cứu", "2.2.Khách thể nghiên cứu"):
            idx_kt = i
            break
    if idx_kt >= 0:
        # heading style ref = "1. Đặt vấn đề"
        ref_idx = find_para(doc, "1. Đặt vấn đề")
        insert_heading_before(doc, idx_kt, "2. Phương pháp nghiên cứu", ref_idx)
        print("[4d] Chèn heading '2. Phương pháp nghiên cứu'")

    # 4e. Renumber tiểu mục: 2.2→2.1, 2.3→2.2, 2.4→2.3
    renumber = [
        ("2.2. Khách thể nghiên cứu", "2.1. Khách thể nghiên cứu"),
        ("2.3. Thiết kế nghiên cứu và công cụ nghiên cứu", "2.2. Thiết kế nghiên cứu và công cụ nghiên cứu"),
        ("2.4. Phân tích dữ liệu", "2.3. Phân tích dữ liệu"),
    ]
    for old, new in renumber:
        idx = find_para(doc, old)
        if idx >= 0:
            replace_para_text(doc.paragraphs[idx], new)
    print("[4e] Đánh số lại tiểu mục 2.1/2.2/2.3")

    # ---- BƯỚC 6: Khai báo AI trong mục 2.3 Phân tích dữ liệu ----
    # Chèn 1 đoạn sau đoạn nội dung của 2.3
    idx_pttdl = find_para_contains(doc, "Dữ liệu được xử lý và phân tích bằng phần mềm")
    if idx_pttdl >= 0:
        ai_text = (
            "Khai báo sử dụng công cụ trí tuệ nhân tạo: trong quá trình hoàn thiện bản thảo, "
            "tác giả có sử dụng công cụ trí tuệ nhân tạo dựa trên mô hình ngôn ngữ lớn (large "
            "language model) ở vai trò hỗ trợ kỹ thuật — rà soát hành văn, kiểm tra định dạng "
            "trích dẫn theo APA phiên bản thứ bảy và sắp xếp tài liệu tham khảo. Toàn bộ thiết "
            "kế nghiên cứu, thu thập và phân tích dữ liệu, diễn giải kết quả và các kết luận "
            "khoa học do tác giả trực tiếp thực hiện và chịu trách nhiệm. Mọi số liệu trình bày "
            "trong bài đều được tác giả đối chiếu với kết quả phân tích gốc trước khi đưa vào "
            "bản thảo. Công cụ trí tuệ nhân tạo không được ghi nhận là tác giả của bài viết."
        )
        insert_para_before(doc, idx_pttdl + 1, ai_text)
        print("[6] Chèn khai báo AI vào mục 2.3 Phân tích dữ liệu")

    # ---- BƯỚC 5: Tuyên bố xung đột lợi ích (trước Lời cảm ơn hoặc TLTK) ----
    idx_loicamon = find_para_contains(doc, "Lời cảm ơn:")
    if idx_loicamon >= 0:
        conflict_text = (
            "Tuyên bố xung đột lợi ích: Tác giả tuyên bố không có bất kỳ xung đột lợi ích nào "
            "liên quan đến nghiên cứu và công bố bài viết này. Tác giả chịu hoàn toàn trách "
            "nhiệm về mọi nhận định, đánh giá và kết luận được trình bày trong bản thảo, cũng "
            "như cam kết nghiên cứu được thực hiện tuân thủ các nguyên tắc đạo đức nghiên cứu "
            "và liêm chính học thuật."
        )
        insert_para_before(doc, idx_loicamon, conflict_text)
        print("[5] Chèn tuyên bố xung đột lợi ích trước Lời cảm ơn")

    # ---- BƯỚC 1b: Sửa tiêu đề Bảng 8 và Bảng 9 trùng nhau ----
    idx_b8 = -1
    idx_b9 = -1
    for i, p in enumerate(doc.paragraphs):
        t = p.text.strip()
        if t.startswith("Bảng 8."):
            idx_b8 = i
        elif t.startswith("Bảng 9."):
            idx_b9 = i
    if idx_b8 >= 0:
        replace_para_text(doc.paragraphs[idx_b8],
                          "Bảng 8. Tóm tắt mô hình hồi quy đa biến dự báo rối loạn lo âu tổng quát")
    if idx_b9 >= 0:
        replace_para_text(doc.paragraphs[idx_b9],
                          "Bảng 9. Hệ số hồi quy trong mô hình đa biến dự báo rối loạn lo âu tổng quát")
    print("[1b] Sửa tiêu đề Bảng 8 và Bảng 9 (trước đây trùng nhau)")

    # ---- BƯỚC 1c: APA → tên đầy đủ lần đầu trong thân bài ----
    # Lần đầu (APA, 2013) trong đoạn "Rối loạn lo âu được đặc trưng..."
    idx_apa1 = find_para_contains(doc, "thường không gắn với một tác nhân cụ thể (APA, 2013)")
    if idx_apa1 >= 0:
        p = doc.paragraphs[idx_apa1]
        new = p.text.replace(
            "(APA, 2013)",
            "(Hiệp hội Tâm thần học Hoa Kỳ [American Psychiatric Association], 2013)", 1)
        replace_para_text(p, new)
        print("[1c] APA → tên đầy đủ (lần đầu xuất hiện)")

    # ---- Metadata cleanup ----
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.subject = ""
    cp.category = ""
    cp.title = ""
    cp.keywords = ""
    cp.created = datetime(2026, 3, 20, 9, 0, 0)
    cp.modified = datetime(2026, 5, 14, 16, 0, 0)

    doc.save(OUT_PATH)
    print(f"\n[DONE] Saved: {OUT_PATH}")


if __name__ == "__main__":
    main()
