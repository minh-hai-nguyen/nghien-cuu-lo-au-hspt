# -*- coding: utf-8 -*-
"""Dump VN018 tom-tat to find exact author string."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
F = os.path.join(ROOT, 'Tom-tat-tung-bai/VN018_AnGiang_2025_YHVN.docx')
d = Document(F)
print(f"=== VN018 tom-tat all paragraphs ===")
for i, p in enumerate(d.paragraphs):
    if p.text.strip():
        print(f"  [{i}] {p.text}")
print(f"\n=== Tables ===")
for ti, tb in enumerate(d.tables):
    print(f"-- Table {ti+1} --")
    for ri, row in enumerate(tb.rows):
        for ci, cell in enumerate(row.cells):
            if cell.text.strip():
                print(f"  r{ri}c{ci}: {cell.text[:300]}")
