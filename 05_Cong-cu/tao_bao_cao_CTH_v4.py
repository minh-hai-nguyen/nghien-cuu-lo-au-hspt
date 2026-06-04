# -*- coding: utf-8 -*-
"""
Báo cáo tổng hợp + Đề cương nghiên cứu
theo phong cách Công Thị Hằng (thuật toán v4)
Bảng Table Grid có viền, căn độ rộng cột chuẩn
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

doc = Document()

# === STYLES ===
style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

# === HELPERS ===
def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_p(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_bp(bold_text, normal_text):
    """Đoạn văn mở đầu bằng text đậm rồi text thường"""
    p = doc.add_paragraph()
    r1 = p.add_run(bold_text)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(normal_text)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return p

def add_i(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_col_width(cell, width_cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(width_cm * 567)))
    w.set(qn('w:type'), 'dxa')
    tcW.append(w)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set widths for ALL rows
    for row_idx, row in enumerate(t.rows):
        for ci in range(len(headers)):
            set_col_width(row.cells[ci], widths[ci])
    # Header row
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    # Data rows
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
    return t

def info_tbl(data):
    t = doc.add_table(rows=len(data), cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        set_col_width(row.cells[0], 3.0)
        set_col_width(row.cells[1], 13.0)
    for i, (k, v) in enumerate(data):
        c0 = t.rows[i].cells[0]
        c0.text = k
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade(c0, 'D9E2F3')
        c1 = t.rows[i].cells[1]
        c1.text = v
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
    return t

# ================================================================
# PHẦN 1: BÁO CÁO TỔNG HỢP THEO PHONG CÁCH CTH v4
# ================================================================
title = doc.add_heading('', level=0)
run = title.add_run('BÁO CÁO TỔNG HỢP\nLO ÂU Ở HỌC SINH TRUNG HỌC CƠ SỞ VÀ TRUNG HỌC PHỔ THÔNG\n(Việt Nam – Đông Nam Á – Thế giới, 2024–2026)')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_i('Ngày báo cáo: 27/03/2026. Phong cách viết: theo thuật toán Công Thị Hằng v4.')
doc.add_paragraph()

# ========== BÀI 1 ==========
add_h('Bài 1', 1)
add_h('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', 2)

add_p('Công trình nghiên cứu « Triệu chứng lo âu và chiến lược ứng phó ở học sinh trung học phổ thông tại Việt Nam sau đại dịch COVID-19: Đánh giá phương pháp hỗn hợp » (Anxiety symptoms and coping strategies among high school students in Vietnam after COVID-19 pandemic: a mixed-method evaluation), do Phạm Thị Thu Hoa và cs. (2024) khảo sát trên mẫu 3.910 học sinh trung học phổ thông từ 14 đến 17 tuổi tại 13 trường ở Hà Nội, Việt Nam.')

add_h('Phương pháp nghiên cứu', 2)

add_p('Công trình này sử dụng thang GAD-7 kết hợp với phỏng vấn sâu qua điện thoại. Nói cách khác, công trình này sử dụng phương pháp hỗn hợp.')

add_p('GAD-7 là phiên bản rút gọn 7 mục của thang Rối loạn Lo âu Tổng quát (Generalized Anxiety Disorder Scale) dành cho thanh thiếu niên và người trưởng thành (Spitzer và cộng sự, 2006). Thang đạt độ tin cậy cao trong nghiên cứu này với Cronbach \u03b1 = 0,916.')

add_p('Tổng quan tài liệu của chúng tôi cho thấy phần lớn các nghiên cứu trước đây tại Việt Nam chỉ sử dụng phương pháp định lượng đơn thuần, bỏ qua bối cảnh văn hóa và trải nghiệm chủ quan của học sinh khi đối mặt với lo âu trong giai đoạn hậu COVID-19.')

add_p('Dữ liệu nhân khẩu xã hội bao gồm giới tính, vị trí địa lý, loại trường và khối lớp.')

add_p('Hướng dẫn phỏng vấn sâu của chúng tôi được thiết kế để tìm hiểu một cách mở, với sự linh hoạt cho các câu hỏi thăm dò. Hai mươi học sinh được phỏng vấn qua điện thoại cho đến khi bão hòa dữ liệu. Các chủ đề phỏng vấn bao gồm nguồn gốc lo âu, biểu hiện triệu chứng và các chiến lược ứng phó.')

add_p('Phương pháp hỗn hợp của chúng tôi sử dụng thang sàng lọc tự báo cáo kết hợp với phỏng vấn sâu, được xây dựng dựa trên các nghiên cứu trước đây chứng minh tính bổ sung của nhiều loại bằng chứng nghiên cứu khác nhau soi chiếu lẫn nhau trên các lĩnh vực tìm hiểu. Ví dụ, trong khi học sinh có thể báo cáo mức lo âu nhẹ trên thang GAD-7, họ có thể kể lại trong phỏng vấn những trải nghiệm lo lắng nghiêm trọng về kỳ vọng gia đình và áp lực thi cử. Ngược lại, một số người tham gia nghiên cứu có thể xác nhận các triệu chứng lo âu trên thang đo mà họ có thể không muốn nói đến trong phỏng vấn tường thuật.')

add_h('Kết quả nghiên cứu định lượng', 2)
add_h('Tỷ lệ và mức độ học sinh có triệu chứng lo âu, theo thang đo GAD-7', 3)

add_p('Về tần suất, mức "nhẹ" là phổ biến nhất : lo âu nhẹ là 23,9% và lo âu trung bình là 10,9%. Mức nặng ít được quan sát thấy hơn (5,8%). Khi gộp lại bất kỳ mức độ triệu chứng nào, khoảng hai phần năm (40,6%) học sinh trong nghiên cứu rơi vào phạm vi từ nhẹ đến nặng của thang sàng lọc lo âu GAD-7.')

tbl(['Mức độ', 'Tỷ lệ'],
    [['Nhẹ', '23,9%'],
     ['Trung bình', '10,9%'],
     ['Nặng', '5,8%'],
     ['Tổng (bất kỳ mức nào)', '40,6%']],
    widths=[5.0, 5.0])
doc.add_paragraph()

add_p('Khác biệt giới tính : Nữ giới luôn báo cáo mức lo âu cao hơn nam giới trong các nghiên cứu quốc tế (McLean và cộng sự, 2011). Thực tế, sự chênh lệch giới trong lo âu ở thanh thiếu niên châu Á đã được ghi nhận ngày càng rõ ràng (Daly, 2022). Chúng tôi xác nhận trong mẫu học sinh THPT Hà Nội rằng điểm GAD-7 trung bình cao hơn ở nữ so với nam (1,74 so với 1,50, p < 0,01). Tương tự, học sinh lớp 12 có điểm lo âu cao hơn lớp 10 và lớp 11 (p < 0,01), và học sinh đô thị cao hơn nông thôn (p < 0,01).')

add_h('Nhận xét, phát hiện qua kết quả nghiên cứu', 2)
add_h('Phát hiện từ phỏng vấn về nguồn gốc lo âu và chiến lược ứng phó', 3)

add_p('*Áp lực học tập và thi cử.* Nổi bật đối với cả nữ sinh và nam sinh là sự lo lắng về kết quả học tập và kỳ thi quan trọng, đặc biệt ở học sinh lớp 12 đối mặt với kỳ thi tốt nghiệp và tuyển sinh đại học. Đây là một lĩnh vực sức khỏe tâm thần không nằm trong các mục sàng lọc của thang GAD-7 nhưng được mô tả chi tiết bởi các cuộc phỏng vấn.')

add_p('*Kỳ thị và mối quan hệ xã hội.* Nỗi lo lắng trước nhận thức của học sinh về sự phán xét từ bạn bè và thái độ kỳ thị trong nhóm xã hội là rất nổi bật. Những vấn đề này được học sinh trải nghiệm như những rào cản đối với sự hòa nhập và tham gia giáo dục.')

add_p('*Kỳ vọng gia đình.* Các cuộc phỏng vấn cho thấy các bối cảnh liên quan đến trải nghiệm lo âu của cả nữ sinh và nam sinh, bao gồm áp lực đáp ứng kỳ vọng học tập của cha mẹ trong cả môi trường trường học và gia đình.')

add_h('Kết luận', 2)

add_p('Dữ liệu của chúng tôi, cho thấy sự hiện diện của triệu chứng lo âu từ nhẹ đến nặng ở khoảng hai phần năm (40,6%) học sinh trung học phổ thông tại Hà Nội trong giai đoạn hậu COVID-19, gợi ý rằng cần chú ý đến các chương trình sàng lọc và can thiệp sớm tại trường học. Sự khác biệt giới tính và địa lý có ý nghĩa, trước đây được tìm thấy trong các nghiên cứu quốc tế, đã có thể quan sát được trong nghiên cứu này. Trong bối cảnh trường trung học phổ thông như nghiên cứu hiện tại, nhu cầu giải quyết mối quan hệ giữa áp lực học tập, kỳ vọng gia đình và sức khỏe tâm thần, dù ở mức nhẹ hay nặng, là thiết yếu để tạo điều kiện cho sự tham gia giáo dục.')

add_h('Tài liệu tham khảo — Bài 1', 3)
add_p('Daly, M. (2022). Prevalence of depression among adolescents in the U.S. from 2009 to 2019: Analysis of trends by sex, race/ethnicity, and income. Journal of Adolescent Health, 70(3), 496–499.')
add_p('McLean, C. P., Asnaani, A., Litz, B. T., & Hofmann, S. G. (2011). Gender differences in anxiety disorders: Prevalence, course of illness, comorbidity and burden of illness. Journal of Psychiatric Research, 45(8), 1027–1035.')
add_p('Phạm Thị Thu Hoa, Đỗ Thị Trang, Nguyễn Thị Liên, & Ngô Anh Vinh. (2024). Anxiety symptoms and coping strategies among high school students in Vietnam after COVID-19 pandemic: A mixed-method evaluation. Frontiers in Public Health, 12, 1232856.')
add_p('Spitzer, R. L., Kroenke, K., Williams, J. B. W., & Löwe, B. (2006). A brief measure for assessing generalized anxiety disorder: The GAD-7. Archives of Internal Medicine, 166(10), 1092–1097.')

doc.add_paragraph()

# ========== BÀI 2 ==========
add_h('Bài 2', 1)
add_h('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', 2)

add_p('Công trình nghiên cứu « Sức khỏe tâm thần ở thanh thiếu niên dân tộc thiểu số tại Việt Nam và các yếu tố tương quan: Nghiên cứu cắt ngang » (Mental health among ethnic minority adolescents in Vietnam and correlated factors: A cross-sectional study), do Ngô Anh Vinh và cs. (2024) khảo sát trên mẫu 845 học sinh nội trú dân tộc thiểu số tại các trường nội trú ở Lạng Sơn, Việt Nam.')

add_h('Phương pháp nghiên cứu', 2)

add_p('Công trình này sử dụng thang DASS-21 kết hợp với Bảng hỏi Trải nghiệm Bất lợi Thời Thơ ấu (ACEs). Nói cách khác, công trình này đánh giá đồng thời sức khỏe tâm thần và tiền sử sang chấn.')

add_p('DASS-21 là phiên bản rút gọn 21 mục của Thang Trầm cảm, Lo âu và Căng thẳng (Depression, Anxiety and Stress Scale) (Lovibond & Lovibond, 1995).')

add_p('ACEs là phiên bản sửa đổi của Bảng hỏi Trải nghiệm Bất lợi Thời Thơ ấu (Adverse Childhood Experiences Questionnaire) (Felitti và cộng sự, 1998).')

add_p('Dữ liệu nhân khẩu xã hội bao gồm tuổi, giới tính, dân tộc, tần suất về thăm nhà, mô hình liên lạc với cha mẹ, thành phần hộ gia đình, chất lượng tình bạn, hoạt động thể chất và kiểu sử dụng internet.')

add_h('Kết quả nghiên cứu định lượng', 2)
add_h('Tỷ lệ và mức độ các vấn đề sức khỏe tâm thần, theo thang đo DASS-21', 3)

add_p('Về tần suất, trầm cảm là phổ biến nhất : gần ba phần năm (59,0%) học sinh có triệu chứng trầm cảm. Lo âu ít hơn nhưng vẫn ở mức rất cao, với hơn một nửa (54,4%) mẫu nghiên cứu. Căng thẳng được báo cáo ở khoảng một phần tư (24,7%) học sinh.')

tbl(['Chỉ số', 'Tỷ lệ'],
    [['Trầm cảm', '59,0%'],
     ['Lo âu', '54,4%'],
     ['Căng thẳng', '24,7%']],
    widths=[5.0, 5.0])
doc.add_paragraph()

add_p('Trải nghiệm bất lợi thời thơ ấu : Gần một nửa (48,9%) học sinh từng trải qua ít nhất một ACE, với số ACE trung bình là 1,1 (SD = 1,8). Số lượng ACEs có tương quan dương với điểm lo âu (hệ số = 0,28), căng thẳng (hệ số = 0,28) và trầm cảm (hệ số = 0,16).')

add_p('Yếu tố xã hội : Tình bạn kém có liên quan đáng kể đến gia tăng lo âu, căng thẳng và trầm cảm. Sử dụng internet hàng ngày hoặc không sử dụng có điểm lo âu thấp hơn so với nhóm chỉ dùng cuối tuần.')

add_h('Nhận xét, phát hiện qua kết quả nghiên cứu', 2)
add_h('Phát hiện về trải nghiệm bất lợi và yếu tố bảo vệ ở học sinh dân tộc thiểu số', 3)

add_p('*Trải nghiệm bất lợi thời thơ ấu.* Nổi bật đối với nhóm học sinh nội trú dân tộc thiểu số là tỷ lệ rất cao trải nghiệm bất lợi — gần một nửa (48,9%) từng trải qua ít nhất một ACE. Đây là một lĩnh vực sức khỏe tâm thần ít được nghiên cứu ở quần thể dân tộc thiểu số Việt Nam nhưng có mối liên hệ rõ ràng với lo âu và trầm cảm hiện tại.')

add_p('*Vai trò của tình bạn.* Chất lượng tình bạn kém có liên quan đáng kể đến gia tăng cả ba chỉ số sức khỏe tâm thần. Những vấn đề này được học sinh trải nghiệm như những rào cản đối với sự hòa nhập trong môi trường nội trú xa gia đình.')

add_p('*Kiểu sử dụng internet.* Các phân tích cho thấy bối cảnh liên quan đến trải nghiệm lo âu bao gồm kiểu sử dụng internet — học sinh dùng internet hàng ngày hoặc không dùng có điểm lo âu thấp hơn so với nhóm chỉ dùng cuối tuần, gợi ý rằng sử dụng không đều đặn có thể phản ánh một kiểu tiếp cận công nghệ bất lợi trong bối cảnh vùng sâu vùng xa.')

add_h('Kết luận', 2)

add_p('Dữ liệu của chúng tôi, cho thấy sự hiện diện của tỷ lệ rất cao các vấn đề sức khỏe tâm thần và trải nghiệm bất lợi thời thơ ấu trong một mẫu học sinh nội trú dân tộc thiểu số, gợi ý rằng cần can thiệp khẩn cấp để giải quyết đồng thời ACEs và các vấn đề nhận thức nhằm nâng cao sức khỏe tâm thần cho nhóm dân số dễ bị tổn thương này. Trong bối cảnh trường nội trú dân tộc thiểu số như nghiên cứu hiện tại, nhu cầu giải quyết mối quan hệ giữa tiền sử sang chấn và sức khỏe tâm thần hiện tại, dù ở mức nhẹ hay nặng, là thiết yếu để tạo điều kiện cho sự tham gia giáo dục.')

add_h('Tài liệu tham khảo — Bài 2', 3)
add_p('Felitti, V. J., Anda, R. F., Nordenberg, D., Williamson, D. F., Spitz, A. M., Edwards, V., ... & Marks, J. S. (1998). Relationship of childhood abuse and household dysfunction to many of the leading causes of death in adults. American Journal of Preventive Medicine, 14(4), 245–258.')
add_p('Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Sydney: Psychology Foundation of Australia.')
add_p('Ngo Anh Vinh, Vu Thi My Hanh, Do Thi Bich Van, Duong Anh Tai, & Le Thi Thanh Thuy. (2024). Mental health among ethnic minority adolescents in Vietnam and correlated factors: A cross-sectional study. Journal of Affective Disorders Reports, 17.')

doc.add_paragraph()

# ========== BÀI 3 ==========
add_h('Bài 3', 1)
add_h('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', 2)

add_p('Công trình nghiên cứu « Khảo sát Sức khỏe Tâm thần Vị thành niên Việt Nam » (Viet Nam Adolescent Mental Health Survey — V-NAMHS), do Vũ Mạnh Lợi và cs. (2022) tại Viện Xã hội học, phối hợp với Đại học Queensland (Úc) và Đại học Johns Hopkins (Mỹ), khảo sát trên mẫu đại diện quốc gia gồm 5.996 vị thành niên từ 10 đến 17 tuổi tại 38 tỉnh thành trên toàn lãnh thổ Việt Nam.')

add_h('Phương pháp nghiên cứu', 2)

add_p('Công trình này sử dụng công cụ chẩn đoán DISC-5 kết hợp với các thang sàng lọc. Nói cách khác, công trình này sử dụng phương pháp chẩn đoán lâm sàng chuẩn hóa, không chỉ sàng lọc triệu chứng.')

add_p('DISC-5 là phiên bản thứ năm của Bảng Phỏng vấn Chẩn đoán dành cho Trẻ em (Diagnostic Interview Schedule for Children, Version 5), được phát triển bởi Đại học Columbia với sự hỗ trợ của Viện Sức khỏe Tâm thần Quốc gia Hoa Kỳ (Bitsko và cộng sự, 2019; Shaffer và cộng sự, 2000). DISC-5 được thiết kế để phỏng vấn viên đã qua đào tạo (không cần là bác sĩ lâm sàng) thực hiện.')

add_p('Tổng quan tài liệu của chúng tôi cho thấy các nghiên cứu trước đây tại Việt Nam bị hạn chế bởi cỡ mẫu nhỏ, phạm vi địa lý hẹp, và thiếu công cụ chẩn đoán — phần lớn chỉ dùng thang sàng lọc triệu chứng như SDQ hay DASS-21, cho tỷ lệ cao hơn nhiều so với công cụ chẩn đoán (Samuels và cộng sự, 2018; Weiss và cộng sự, 2014).')

add_h('Kết quả nghiên cứu', 2)
add_h('Tỷ lệ vấn đề sức khỏe tâm thần và rối loạn tâm thần, theo DISC-5 và DSM-5', 3)

add_p('Về vấn đề sức khỏe tâm thần (mental health problems — bao gồm triệu chứng dưới ngưỡng và đủ ngưỡng), khoảng một phần năm (21,7%) vị thành niên báo cáo ít nhất một vấn đề SKTT trong 12 tháng qua. Lo âu là phổ biến nhất : 18,6% có vấn đề lo âu, tiếp theo là trầm cảm 4,3%, vấn đề chú ý/tăng động 2,8%, stress sau sang chấn 1,0% và rối loạn hành vi 0,7%.')

tbl(['Loại vấn đề SKTT', 'Tỷ lệ (vấn đề SKTT)', 'Tỷ lệ (rối loạn, DSM-5)'],
    [['Lo âu (ám sợ xã hội + RLLA tổng quát)', '18,6%', '2,3%'],
     ['Trầm cảm', '4,3%', '0,9%'],
     ['Stress sau sang chấn', '1,0%', '—'],
     ['Rối loạn hành vi', '0,7%', '—'],
     ['Vấn đề chú ý/tăng động', '2,8%', '—'],
     ['Bất kỳ vấn đề nào', '21,7%', '3,3%']],
    widths=[5.5, 3.5, 3.5])
doc.add_paragraph()

add_p('Phát hiện then chốt : Tỷ lệ rối loạn lo âu theo chẩn đoán DSM-5 chỉ là 2,3%, và tỷ lệ bất kỳ rối loạn tâm thần nào là 3,3% — thấp hơn rất nhiều so với tỷ lệ sàng lọc DASS-21 (25–86%) trong các nghiên cứu khác tại Việt Nam. Không có sự khác biệt có ý nghĩa giữa nam và nữ, hoặc giữa nhóm 10–13 tuổi và 14–17 tuổi.')

add_h('Kết luận', 2)

add_p('Dữ liệu của chúng tôi, cho thấy sự hiện diện của các vấn đề sức khỏe tâm thần ở khoảng một phần năm (21,7%) vị thành niên Việt Nam trên mẫu đại diện quốc gia, trong đó lo âu là phổ biến nhất, gợi ý rằng cần chú ý đến khoảng cách lớn giữa tỷ lệ sàng lọc và tỷ lệ chẩn đoán khi hoạch định chính sách. Sự khác biệt đáng kể giữa tỷ lệ sàng lọc (18,6% vấn đề lo âu) và tỷ lệ chẩn đoán (2,3% rối loạn lo âu) có ý nghĩa lâm sàng tiềm năng cho việc thiết kế chương trình can thiệp phù hợp — cả cho nhóm đủ ngưỡng chẩn đoán lẫn nhóm có triệu chứng dưới ngưỡng hội chứng.')

add_h('Tài liệu tham khảo — Bài 3', 3)
add_p('American Psychiatric Association. (2013). Diagnostic and Statistical Manual of Mental Disorders (5th ed.). Washington, DC: APA.')
add_p('Bitsko, R. H., Holbrook, J. R., Ghandour, R. M., et al. (2019). Epidemiology and impact of health care provider-diagnosed anxiety and depression among US children. Journal of Developmental & Behavioral Pediatrics, 39(5), 395–403.')
add_p('Samuels, F., Jones, N., & Gupta, T. (2018). Mental health and psychosocial well-being of adolescents in Viet Nam: Findings from a mixed-method baseline study. London: Overseas Development Institute.')
add_p('Shaffer, D., Fisher, P., Lucas, C. P., et al. (2000). NIMH Diagnostic Interview Schedule for Children Version IV (NIMH DISC-IV). Journal of the American Academy of Child & Adolescent Psychiatry, 39(1), 28–38.')
add_p('Vu Manh Loi, Nguyen Duc Vinh, Dao Thi Khanh Hoa, et al. (2022). Viet Nam Adolescent Mental Health Survey (V-NAMHS): Report on Main Findings. Hanoi: Institute of Sociology.')
add_p('Weiss, B., Dang, M., Trung, L., Nguyen, M. C., Thuy, N. T. H., & Pollack, A. (2014). A nationally-representative epidemiological and risk factor assessment of child mental health in Vietnam. International Perspectives in Psychology, 3(3), 139–153.')

doc.add_paragraph()

# ========== BÀI 4 (GBD ASEAN) ==========
add_h('Bài 4', 1)
add_h('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', 2)

add_p('Công trình nghiên cứu « Dịch tễ học và gánh nặng mười rối loạn tâm thần tại các quốc gia thuộc Hiệp hội các Quốc gia Đông Nam Á (ASEAN), 1990–2021: Phát hiện từ Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2021 » (The epidemiology and burden of ten mental disorders in countries of the Association of Southeast Asian Nations, 1990–2021: findings from the Global Burden of Disease Study 2021), do Nhóm Hợp tác GBD 2021 về Rối loạn Tâm thần ASEAN (2025) phân tích dữ liệu từ 10 quốc gia ASEAN gồm Brunei, Campuchia, Indonesia, Lào, Malaysia, Myanmar, Philippines, Singapore, Thái Lan và Việt Nam.')

add_h('Phương pháp nghiên cứu', 2)

add_p('Công trình này sử dụng mô hình DisMod-MR 2.1 phân tích dữ liệu GBD 2021. Nói cách khác, công trình này phân tích dữ liệu thứ cấp từ nghiên cứu gánh nặng bệnh tật toàn cầu.')

add_p('DisMod-MR 2.1 là công cụ hồi quy meta Bayesian (Bayesian meta-regression modelling tool) được phát triển bởi Viện Đo lường và Đánh giá Sức khỏe (IHME) để ước tính tỷ lệ hiện mắc theo tuổi, giới, năm và địa điểm.')

add_h('Kết quả nghiên cứu', 2)

add_p('Về tỷ lệ rối loạn tâm thần chung, tỷ lệ chuẩn hóa theo tuổi tại ASEAN là 11,9% (95% UI: 10,9–12,9) năm 2021. Lo âu và trầm cảm là phổ biến nhất. Dao động từ 10,1% tại Việt Nam đến 13,2% tại Malaysia.')

tbl(['Chỉ số', 'Giá trị (95% UI)'],
    [['Tổng số ca RLTT (2021)', '80,4 triệu (73,8–87,2)'],
     ['Tăng so với 1990', '70,0% (63,5–77,2)'],
     ['Tỷ lệ chuẩn hóa theo tuổi', '11,9% (10,9–12,9)'],
     ['Tỷ lệ — Việt Nam', '10,1% (9,1–11,3)'],
     ['Tỷ lệ — Malaysia (cao nhất)', '13,2% (11,6–15,3)'],
     ['DALYs (2021)', '11,2 triệu (8,5–14,3)'],
     ['DALYs tăng so với 1990', '87,4% (81,1–94,0)']],
    widths=[6.0, 8.0])
doc.add_paragraph()

add_p('Gánh nặng theo nhóm tuổi : Nhóm 10–14 tuổi có gánh nặng bệnh tật do rối loạn tâm thần cao nhất — chiếm 16,3% (12,7–20,5) tổng DALYs ở nhóm tuổi này. Rối loạn tâm thần nằm trong 10 nguyên nhân gánh nặng bệnh tật hàng đầu tại tất cả các nước ASEAN trừ Myanmar.')

add_h('Kết luận', 2)

add_p('Dữ liệu của chúng tôi, cho thấy sự gia tăng 70% số ca rối loạn tâm thần tại ASEAN từ 1990 đến 2021 với nhóm 10–14 tuổi chịu gánh nặng lớn nhất, gợi ý rằng cần một cách tiếp cận liên ngành toàn diện để giải quyết khoảng trống phòng ngừa và điều trị trên toàn bộ quần thể. Sự gia tăng có thể phản ánh một phần sự cải thiện trong phát hiện bệnh, nhưng gánh nặng thực sự vẫn rất lớn và chưa được đáp ứng đầy đủ.')

add_h('Tài liệu tham khảo — Bài 4', 3)
add_p('GBD 2021 ASEAN Mental Disorders Collaborators. (2025). The epidemiology and burden of ten mental disorders in countries of the Association of Southeast Asian Nations (ASEAN), 1990–2021: Findings from the Global Burden of Disease Study 2021. The Lancet Public Health, 10, e480.')

doc.add_paragraph()

# ========== BÀI 5 (Hoàng Trung Học — so sánh trong/sau COVID) ==========
add_h('Bài 5', 1)
add_h('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', 2)

add_p('Công trình nghiên cứu « Mức độ căng thẳng, lo âu và trầm cảm ở vị thành niên trong và sau đại dịch COVID-19 tại Việt Nam: Nghiên cứu cắt ngang » (Levels of Stress, Anxiety, and Depression in Adolescents during and after the COVID-19 Pandemic in Vietnam: A Cross-sectional study), do Hoàng Trung Học và Nguyễn Thùy Dung (2025) khảo sát trên mẫu 8.473 vị thành niên tại 6 tỉnh thành Việt Nam.')

add_h('Phương pháp nghiên cứu', 2)
add_p('Công trình này sử dụng thang DASS-21 phiên bản tiếng Việt. Nói cách khác, công trình này sử dụng phương pháp sàng lọc triệu chứng cắt ngang, so sánh hai giai đoạn.')

add_h('Kết quả nghiên cứu định lượng', 2)

tbl(['Chỉ số', 'Trong đại dịch', 'Sau đại dịch', 'Thay đổi'],
    [['Căng thẳng', '65,5%', '55,4%', '\u2193 10,1%'],
     ['Lo âu', '41,5%', '25,4%', '\u2193 16,1%'],
     ['Trầm cảm', '34,2%', '20,1%', '\u2193 14,1%']],
    widths=[4.0, 3.5, 3.5, 3.0])
doc.add_paragraph()

add_p('Khác biệt giai đoạn : Các nghiên cứu quốc tế luôn báo cáo sự gia tăng các vấn đề sức khỏe tâm thần ở vị thành niên trong đại dịch COVID-19 (Panchal và cộng sự, 2023). Thực tế, một số quốc gia đã ghi nhận sự hồi phục một phần sau đại dịch nhưng không về mức trước đại dịch (Daly, 2022). Chúng tôi xác nhận trong mẫu vị thành niên Việt Nam rằng tất cả các chỉ số đều giảm sau đại dịch, trong đó lo âu giảm nhiều nhất (16,1 điểm phần trăm), nhưng vẫn ở mức đáng lo ngại.')

add_p('Yếu tố nguy cơ : Các yếu tố liên quan đến mức lo âu cao bao gồm chất lượng mối quan hệ cha mẹ–con kém, sử dụng thiết bị điện tử quá mức và chất lượng giấc ngủ kém.')

add_h('Nhận xét, phát hiện qua kết quả nghiên cứu', 2)
add_h('Phát hiện về yếu tố bảo vệ và xu hướng phục hồi', 3)

add_p('*Vai trò bảo vệ của gia đình.* Nổi bật đối với giai đoạn sau đại dịch là vai trò bảo vệ của mối quan hệ gia đình tích cực, đặc biệt trong việc giảm lo âu. Đây là một yếu tố không nằm trong các mục sàng lọc của thang DASS-21 nhưng được xác nhận qua phân tích đa biến.')

add_p('*Hoạt động ngoại khóa.* Sự tham gia các hoạt động ngoại khóa và lối sống lành mạnh có tác dụng bảo vệ rõ rệt hơn trong giai đoạn sau đại dịch so với trong đại dịch. Những phát hiện này được trải nghiệm trong bối cảnh mở cửa trường học trở lại, cho phép tái thiết lập các mối quan hệ xã hội trực tiếp.')

add_h('Kết luận', 2)
add_p('Dữ liệu của chúng tôi, cho thấy sự giảm đáng kể nhưng vẫn ở mức cao của các triệu chứng lo âu, trầm cảm và căng thẳng ở vị thành niên Việt Nam sau đại dịch trong một mẫu lớn gồm 8.473 người tham gia tại 6 tỉnh thành, gợi ý rằng việc điều chỉnh sử dụng công nghệ và tăng cường sự tham gia của gia đình là thiết yếu cho các chiến lược phòng ngừa và can thiệp sức khỏe tâm thần trong giai đoạn khủng hoảng và phục hồi. Các yếu tố bảo vệ đã xác định, trước đây được tìm thấy trong các nghiên cứu phương Tây, đã có thể quan sát được trong bối cảnh Việt Nam.')

add_h('Tài liệu tham khảo — Bài 5', 3)
add_p('Daly, M. (2022). Prevalence of depression among adolescents in the U.S. from 2009 to 2019. Journal of Adolescent Health, 70(3), 496–499.')
add_p('Hoang Trung Hoc & Nguyen Thuy Dung. (2025). Levels of Stress, Anxiety, and Depression in Adolescents during and after the COVID-19 Pandemic in Vietnam: A Cross-sectional study. American Journal of Psychiatric Rehabilitation, 28(1).')
add_p('Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Sydney: Psychology Foundation of Australia.')
add_p('Panchal, U., Salazar de Pablo, G., Franco, M., et al. (2023). The impact of COVID-19 lockdown on child and adolescent mental health: Systematic review. European Child & Adolescent Psychiatry, 32, 1151–1177.')

doc.add_paragraph()

# ================================================================
# PHẦN 2: ĐỀ CƯƠNG NGHIÊN CỨU
# ================================================================
doc.add_page_break()

title2 = doc.add_heading('', level=0)
run2 = title2.add_run('ĐỀ CƯƠNG NGHIÊN CỨU\nLO ÂU Ở HỌC SINH TRUNG HỌC CƠ SỞ VÀ TRUNG HỌC PHỔ THÔNG TẠI VIỆT NAM:\nTỶ LỆ, YẾU TỐ LIÊN QUAN VÀ HIỆU QUẢ CAN THIỆP TẠI TRƯỜNG HỌC')
run2.font.name = 'Times New Roman'
run2.font.size = Pt(14)
run2.font.color.rgb = RGBColor(0, 51, 102)
title2.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_i('Đề cương theo phong cách Công Thị Hằng v4. Dựa trên tổng hợp 21 bài nghiên cứu 2024–2026.')
doc.add_paragraph()

add_h('1. Đặt vấn đề', 1)

add_p('Rối loạn lo âu ở vị thành niên đã trở thành mối quan tâm sức khỏe công cộng toàn cầu trong thập kỷ qua. Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2021 cho thấy nhóm 10–14 tuổi tại các nước ASEAN chịu gánh nặng bệnh tật do rối loạn tâm thần cao nhất, chiếm 16,3% tổng DALYs ở nhóm tuổi này (GBD 2021 ASEAN Mental Disorders Collaborators, 2025). Thực tế, tỷ lệ rối loạn tâm thần tại ASEAN đã tăng 70% từ năm 1990 đến 2021, với 80,4 triệu ca được ghi nhận.')

add_p('Tại Việt Nam, tổng quan tài liệu của chúng tôi cho thấy các ước tính tỷ lệ lo âu ở học sinh THCS và THPT dao động rất lớn — từ 2,3% khi dùng công cụ chẩn đoán DISC-5 (Vũ Mạnh Lợi và cộng sự, 2022) đến 86,2% khi dùng thang sàng lọc DASS-21 (Nguyễn Ngọc Bảo Quyên và cộng sự, 2025). Sự chênh lệch này phản ánh khoảng trống phương pháp luận nghiêm trọng: thiếu nghiên cứu dùng công cụ chẩn đoán chuẩn hóa, thiếu nghiên cứu dọc, và thiếu dữ liệu đại diện quốc gia.')

add_p('Đặc biệt, ba khoảng trống chính được nhận diện. Thứ nhất, chưa có thử nghiệm ngẫu nhiên có đối chứng (RCT) nào về can thiệp sức khỏe tâm thần tại trường học Việt Nam — trong khi tổng quan hệ thống của Zhameden và cs. (2025) chỉ tìm được 6 RCTs tại tất cả các nước thu nhập thấp và trung bình, không có bài nào từ Việt Nam. Thứ hai, nhóm dân tộc thiểu số (với tỷ lệ lo âu 54,4% theo Ngô Anh Vinh và cộng sự, 2024) và nhóm nông thôn hầu như chưa được nghiên cứu. Thứ ba, vai trò trung gian của giấc ngủ, mạng xã hội và áp lực học thêm chưa được đánh giá trong mô hình nhân quả.')

add_h('2. Mục tiêu nghiên cứu', 1)

add_p('Mục tiêu tổng quát: Xác định tỷ lệ, yếu tố liên quan và đánh giá hiệu quả can thiệp tại trường học đối với lo âu ở học sinh THCS và THPT tại Việt Nam.')

add_p('Mục tiêu cụ thể:')
add_p('(1) Xác định tỷ lệ rối loạn lo âu (chẩn đoán) và triệu chứng lo âu (sàng lọc) ở học sinh THCS và THPT tại 3 vùng sinh thái (đồng bằng, miền núi, đô thị).')
add_p('(2) Phân tích các yếu tố nguy cơ và bảo vệ: áp lực học tập, mối quan hệ gia đình, sử dụng mạng xã hội, chất lượng giấc ngủ, trải nghiệm bất lợi thời thơ ấu (ACEs), và hỗ trợ xã hội.')
add_p('(3) Thiết kế, triển khai và đánh giá hiệu quả một chương trình can thiệp sức khỏe tâm thần tại trường học (school-based intervention) sử dụng phương pháp CBT nhóm.')

add_h('3. Phương pháp nghiên cứu', 1)

add_p('Nghiên cứu của chúng tôi sử dụng thiết kế hỗn hợp tuần tự giải thích (sequential explanatory mixed-methods design) gồm ba giai đoạn. Nói cách khác, nghiên cứu kết hợp khảo sát cắt ngang quy mô lớn, thử nghiệm can thiệp và phỏng vấn sâu.')

add_h('3.1 Giai đoạn 1 — Khảo sát cắt ngang (năm 1)', 2)

info_tbl([
    ('Thiết kế', 'Nghiên cứu cắt ngang đa tầng (multi-level cross-sectional)'),
    ('Mẫu', '3.000 học sinh THCS + 3.000 học sinh THPT tại 30 trường thuộc 3 vùng (đồng bằng, miền núi, đô thị) ở 6 tỉnh/thành'),
    ('Công cụ sàng lọc', 'GAD-7 (lo âu), PHQ-9A (trầm cảm), PSQI (giấc ngủ), ESSA (căng thẳng học tập), ACEs (trải nghiệm bất lợi)'),
    ('Công cụ chẩn đoán', 'DISC-5 module Lo âu — trên mẫu con 600 HS để so sánh sàng lọc vs chẩn đoán'),
    ('Phân tích', 'Hồi quy logistic đa tầng (multilevel logistic regression), phân tích đường dẫn (path analysis) cho mô hình trung gian')
])
doc.add_paragraph()

add_h('3.2 Giai đoạn 2 — Thử nghiệm can thiệp RCT (năm 2)', 2)

info_tbl([
    ('Thiết kế', 'Thử nghiệm ngẫu nhiên có đối chứng theo cụm (cluster-RCT), 2 nhánh'),
    ('Mẫu', '600 HS có triệu chứng lo âu (GAD-7 \u2265 5) tại 20 trường (10 can thiệp + 10 đối chứng)'),
    ('Can thiệp', 'CBT nhóm tại trường, 8 buổi/2 tháng, do giáo viên tư vấn đã được đào tạo thực hiện'),
    ('Đối chứng', 'Chương trình giáo dục sức khỏe thông thường'),
    ('Đo lường', 'GAD-7, PHQ-9A tại 4 thời điểm: T0 (trước), T1 (sau 2 tháng), T2 (6 tháng), T3 (12 tháng)'),
    ('Phân tích', 'Mô hình hỗn hợp tuyến tính (linear mixed model), intention-to-treat, phân tích cost-effectiveness')
])
doc.add_paragraph()

add_h('3.3 Giai đoạn 3 — Phỏng vấn sâu (năm 2–3)', 2)

info_tbl([
    ('Thiết kế', 'Nghiên cứu định tính — phỏng vấn bán cấu trúc'),
    ('Mẫu', '40 HS (20 nhóm can thiệp + 20 đối chứng), 20 phụ huynh, 10 giáo viên tư vấn'),
    ('Phương pháp', 'Phân tích mã hóa chủ đề định tính bằng NVivo'),
    ('Mục đích', 'Tìm hiểu trải nghiệm chủ quan về lo âu, đánh giá tính chấp nhận và khả thi của can thiệp CBT')
])
doc.add_paragraph()

add_h('4. Ý nghĩa khoa học và thực tiễn', 1)

add_p('Nghiên cứu này sẽ là công trình đầu tiên tại Việt Nam sử dụng đồng thời công cụ chẩn đoán DISC-5 và thang sàng lọc GAD-7 trên cùng một mẫu, cho phép xác định ngưỡng cắt phù hợp cho quần thể Việt. Thử nghiệm RCT can thiệp CBT tại trường học sẽ lấp đầy khoảng trống bằng chứng quan trọng mà tổng quan hệ thống của Zhameden và cs. (2025) đã chỉ ra — chưa có RCT nào từ Việt Nam.')

add_p('Về thực tiễn, kết quả sẽ cung cấp dữ liệu cho Bộ Giáo dục và Đào tạo để xây dựng chương trình sức khỏe tâm thần học đường chuẩn hóa, đặc biệt cho nhóm dân tộc thiểu số và vùng nông thôn — nơi tỷ lệ lo âu cao nhưng tiếp cận dịch vụ thấp nhất.')

add_h('5. Tiến độ dự kiến', 1)

tbl(['Giai đoạn', 'Thời gian', 'Hoạt động chính'],
    [['Chuẩn bị', 'Tháng 1–6 năm 1', 'Xin phê duyệt đạo đức, đào tạo nhóm, dịch và thử nghiệm công cụ'],
     ['GĐ1: Cắt ngang', 'Tháng 7–12 năm 1', 'Khảo sát 6.000 HS tại 30 trường, 6 tỉnh'],
     ['Phân tích GĐ1', 'Tháng 1–3 năm 2', 'Phân tích dữ liệu, thiết kế can thiệp'],
     ['GĐ2: RCT', 'Tháng 4 năm 2 – tháng 3 năm 3', 'Can thiệp CBT + theo dõi 12 tháng'],
     ['GĐ3: Định tính', 'Tháng 10 năm 2 – tháng 6 năm 3', 'Phỏng vấn sâu 70 người'],
     ['Tổng hợp', 'Tháng 7–12 năm 3', 'Phân tích tích hợp, viết báo cáo, công bố']],
    widths=[3.0, 4.0, 9.0])
doc.add_paragraph()

add_h('6. Tài liệu tham khảo (trích)', 1)

refs = [
    'GBD 2021 ASEAN Mental Disorders Collaborators (2025). The epidemiology and burden of ten mental disorders in countries of the ASEAN, 1990\u20132021. The Lancet Public Health, 10, e480.',
    'Hoa PTT, Trang DT, Lien NT, Vinh NA (2024). Anxiety symptoms and coping strategies among high school students in Vietnam after COVID-19. Frontiers in Public Health, 12, 1232856.',
    'Hoang TH, Nguyen TD (2025). Levels of Stress, Anxiety, and Depression in Adolescents during and after the COVID-19 Pandemic in Vietnam. Am J Psychiatric Rehabilitation, 28(1).',
    'Ngo AV, Vu TMH, Do TBV, Duong AT, Le TTT (2024). Mental health among ethnic minority adolescents in Vietnam. J Affective Disorders Reports, 17.',
    'Vu ML, Nguyen DV, Dao TKH, et al. (2022). Viet Nam Adolescent Mental Health Survey (V-NAMHS): Report on Main Findings. Hanoi: Institute of Sociology.',
    'Yin SZD, Low MK, Mishu MP (2025). School-based interventions to prevent anxiety and depression in LMICs. PLoS One, 20(4), e0316825.',
    'Anderson TL et al. (2025). Contributing Factors to the Rise in Adolescent Anxiety. J Child Adolesc Psychiatric Nursing, 38, e70009.',
    'Salari N et al. (2024). Global Prevalence of Social Anxiety Disorder in Children, Adolescents and Youth. J Prevention.',
]

for i, ref in enumerate(refs, 1):
    add_p(f'{i}. {ref}')

doc.add_paragraph()
add_i('Tài liệu tạo ngày 27/03/2026. Viết theo thuật toán bắt chước phong cách Công Thị Hằng v4.')

# === SAVE ===
output = r'c:\Users\HLC\OneDrive\read_books\Lo-au\Lo au - Bao cao CTH + De cuong - 27032026.docx'
doc.save(output)
print('Done: ' + output)
