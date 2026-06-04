# -*- coding: utf-8 -*-
"""Find citations in the references list to cite in critique."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

d = Document(PATH)

# Find Tai lieu tham khao section
refs_start = None
for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if 'Tài liệu tham khảo' in t or 'References' in t.lower() or 'TÀI LIỆU THAM KHẢO' in t:
        refs_start = i
        print(f'Refs start at #{i}: {t[:100]}')
        break

# Dump 50 refs
if refs_start:
    for i in range(refs_start, min(refs_start + 60, len(d.paragraphs))):
        t = d.paragraphs[i].text.strip()
        if t and len(t) > 40:
            print(f'  [{i}] {t[:200]}')
