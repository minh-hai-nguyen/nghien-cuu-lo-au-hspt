"""Build doc: Weighted prevalence - Ty le co trong so."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

title = d.add_heading('Weighted Prevalence (Tỷ lệ có trọng số): Là gì và vì sao cần?', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Giải thích chỉ số quan trọng trong khảo sát dịch tễ quy mô lớn')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# META
meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
mp = cell.paragraphs[0]
mp.add_run('Câu hỏi gốc của thầy: ').bold = True
mp.add_run('"Tỷ lệ có trọng số (weighted prevalence) nghĩa là như thế nào?"')
mp2 = cell.add_paragraph()
mp2.add_run('Bối cảnh: ').bold = True
mp2.add_run('Chen et al. (2023) BMC Psychiatry (QT007 — Bài 7) báo cáo "Tỷ lệ có trọng số của triệu chứng trầm cảm và lo âu (N=63.205)" trên mẫu HS trung học Tự Cống, Trung Quốc — lấy mẫu cụm từ 2 quận + 1 huyện.')
d.add_paragraph()

# =========================================================
# PHẦN 1 — PREVALENCE LÀ GÌ
# =========================================================
add_heading('1. Prevalence (tỷ lệ hiện mắc) là gì?', level=1)

add_para('Prevalence = Tỷ lệ người có điều kiện/triệu chứng quan tâm tại một thời điểm, tính trên tổng số người trong quần thể.', bold=True)
d.add_paragraph()

add_para('Công thức thường:', bold=True)
add_para('p = số người có triệu chứng / tổng số người trong mẫu', size=12, color=(192, 0, 0))
d.add_paragraph()

add_para('Ví dụ: khảo sát 1.000 HS, 200 HS có triệu chứng lo âu → prevalence = 200/1.000 = 20 %.')
d.add_paragraph()

add_para('Nhưng trong các khảo sát dịch tễ quy mô lớn (quốc gia, vùng), lấy mẫu KHÔNG ĐƠN GIẢN — thường dùng stratified cluster sampling hoặc multi-stage sampling. Khi đó công thức đơn giản trên KHÔNG CHÍNH XÁC.')
d.add_paragraph()

# =========================================================
# PHẦN 2 — VẤN ĐỀ
# =========================================================
add_heading('2. Vì sao công thức đơn giản không chính xác?', level=1)

add_para('2.1. Vấn đề 1 — Xác suất chọn KHÔNG đều', bold=True, size=12)
add_para('Trong khảo sát thực tế, không phải HS nào cũng có xác suất như nhau để được chọn vào mẫu:')
for it in [
    'Trường LỚN (2.000 HS) được chọn 200 HS → xác suất mỗi HS = 200/2.000 = 10 %',
    'Trường NHỎ (300 HS) được chọn 200 HS → xác suất mỗi HS = 200/300 = 67 %',
    'HS trường nhỏ "over-represented" (được đại diện quá mức) so với HS trường lớn',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

add_para('→ Nếu tính tỷ lệ đơn giản, kết quả bị lệch theo cấu trúc mẫu, không phản ánh đúng quần thể thực tế.')
d.add_paragraph()

add_para('2.2. Vấn đề 2 — Over-sampling theo chủ đích', bold=True, size=12)
add_para('Nhiều khảo sát cố tình lấy mẫu "thừa" ở một số nhóm nhỏ để có đủ số liệu phân tích cho nhóm đó. Ví dụ:')
for it in [
    'Khảo sát toàn quốc: HS dân tộc thiểu số chỉ chiếm 15 %, nhưng tác giả chọn 30 % mẫu là HS DTTS để có đủ power phân tích nhóm này',
    'Nếu dùng công thức đơn giản → tỷ lệ tổng sẽ LỆCH về đặc điểm của nhóm over-sampled',
    'Phải weighting lại để "đẩy" HS DTTS xuống tỷ lệ thực (15 %) khi tính prevalence tổng',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# PHẦN 3 — WEIGHTED PREVALENCE
# =========================================================
add_heading('3. Weighted prevalence — Giải pháp', level=1, color=(192, 80, 77))

add_para('Ý tưởng cốt lõi: Mỗi HS được gán một TRỌNG SỐ (weight), bằng 1 / xác suất HS đó được chọn. Weight nói rằng "HS này đại diện cho bao nhiêu HS khác trong quần thể".', bold=True)
d.add_paragraph()

add_para('Công thức weighted prevalence:', bold=True)
add_para('p = Σ(wᵢ × yᵢ) / Σ(wᵢ)', size=14, bold=True, color=(192, 0, 0))
d.add_paragraph()

add_para('Trong đó:')
for it in [
    'wᵢ = trọng số của HS thứ i (= 1 / xác suất được chọn, có thể điều chỉnh thêm cho non-response)',
    'yᵢ = 1 nếu HS i có triệu chứng, 0 nếu không',
    'Σ(wᵢ) = tổng trọng số = ước tính cỡ quần thể',
    'Σ(wᵢ × yᵢ) = ước tính số người có triệu chứng trong quần thể',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# Box so sánh
add_para('So sánh công thức:', bold=True, size=12)
cmp_tbl = d.add_table(rows=3, cols=2)
cmp_tbl.style = 'Table Grid'
cmps = [
    ('Unweighted (tỷ lệ thường)',
     'p = n_có / n_tổng\nMỗi HS đóng góp NHƯ NHAU'),
    ('Weighted (tỷ lệ có trọng số)',
     'p = Σ(wᵢ × yᵢ) / Σ(wᵢ)\nMỗi HS đóng góp THEO TRỌNG SỐ wᵢ'),
    ('Khi nào 2 công thức cho kết quả GIỐNG NHAU?',
     'Khi tất cả wᵢ = hằng số (mỗi HS có xác suất được chọn như nhau) — tức SRS (simple random sampling). Trong survey phức tạp, 2 công thức thường cho kết quả KHÁC nhau.'),
]
for i, (k, v) in enumerate(cmps):
    row = cmp_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v
d.add_paragraph()

# =========================================================
# PHẦN 4 — VÍ DỤ CỤ THỂ
# =========================================================
add_heading('4. Ví dụ tính weighted prevalence', level=1)

add_para('Giả sử khảo sát 3 trường, chọn 100 HS mỗi trường:', bold=True)
d.add_paragraph()

ex_tbl = d.add_table(rows=5, cols=5)
ex_tbl.style = 'Table Grid'
hdr = ex_tbl.rows[0]
for i, h in enumerate(['Trường', 'Tổng HS', 'HS được chọn', 'HS có lo âu', 'Weight (wᵢ)']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

exs = [
    ('A (lớn)', '2.000', '100', '30', '2.000/100 = 20'),
    ('B (vừa)', '1.000', '100', '40', '1.000/100 = 10'),
    ('C (nhỏ)', '300', '100', '50', '300/100 = 3'),
    ('Tổng', '3.300', '300', '120', ''),
]
for i, row_data in enumerate(exs):
    row = ex_tbl.rows[i+1]
    for j, v in enumerate(row_data):
        row.cells[j].text = v
        if i == 3:
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True
d.add_paragraph()

# Tính 2 cách
add_para('CÁCH 1 — Unweighted (sai):', bold=True, color=(192, 0, 0))
add_para('p = 120 / 300 = 40 %', size=12)
add_para('→ Kết luận "40% HS lo âu" — NHƯNG đây là tỷ lệ trong MẪU, không phản ánh đúng quần thể 3.300 HS.')
d.add_paragraph()

add_para('CÁCH 2 — Weighted (đúng):', bold=True, color=(54, 95, 44))
calc_tbl = d.add_table(rows=4, cols=1)
calc_tbl.style = 'Table Grid'
calcs = [
    'Σ(wᵢ × yᵢ) = 30×20 + 40×10 + 50×3 = 600 + 400 + 150 = 1.150',
    'Σ(wᵢ) = 100×20 + 100×10 + 100×3 = 2.000 + 1.000 + 300 = 3.300',
    'p_weighted = 1.150 / 3.300 ≈ 34,8 %',
    '→ Kết luận "34,8% HS lo âu" — phản ánh ĐÚNG tỷ lệ trong quần thể 3.300 HS thật.',
]
for i, c in enumerate(calcs):
    cell = calc_tbl.rows[i].cells[0]
    shade_cell(cell, 'E2EFDA' if i < 3 else 'FFF2CC')
    p = cell.paragraphs[0]
    r = p.add_run(c)
    if i == 3:
        r.bold = True
    r.font.size = Pt(11)
d.add_paragraph()

add_para('Chênh lệch giữa 2 cách: 40 % (unweighted) vs 34,8 % (weighted) = 5,2 điểm phần trăm. Đáng kể! Đây là lý do không thể "lười" dùng tỷ lệ thường khi thiết kế lấy mẫu phức tạp.', color=(192, 0, 0))
d.add_paragraph()

# =========================================================
# PHẦN 5 — TRỌNG SỐ THEO DÕI NHỮNG GÌ
# =========================================================
add_heading('5. Weight thường kết hợp những thành phần nào?', level=1)

add_para('Trong survey thực tế, weight KHÔNG chỉ là 1/xác suất chọn. Nó thường là TÍCH của nhiều thành phần:', bold=False)
d.add_paragraph()

w_tbl = d.add_table(rows=5, cols=2)
w_tbl.style = 'Table Grid'
ws = [
    ('① Design weight (trọng số thiết kế)',
     'w_design = 1 / xác suất được chọn vào mẫu. Đây là thành phần CƠ BẢN nhất.'),
    ('② Non-response adjustment (điều chỉnh không phản hồi)',
     'Một số HS từ chối hoặc vắng mặt → không có data. Weight của HS còn lại được TĂNG LÊN để bù đắp. Nếu không điều chỉnh, kết quả lệch về những người DỄ phản hồi (có thể khác với người từ chối).'),
    ('③ Post-stratification (hậu phân tầng)',
     'Điều chỉnh sao cho tỷ lệ nam/nữ, độ tuổi, vùng miền trong MẪU có trọng số khớp với tỷ lệ trong QUẦN THỂ thực (lấy từ Tổng điều tra dân số). Giúp kết quả ngoại suy cho toàn quốc chính xác hơn.'),
    ('④ Trimming (cắt xén trọng số cực đoan)',
     'Nếu 1 HS có weight quá lớn (vd 1.000) → biến thành outlier, ảnh hưởng ước tính. Cắt xuống mức tối đa (vd 100) để ổn định SE.'),
    ('Weight cuối cùng',
     'w_final = w_design × w_nonresponse × w_poststrat (sau trimming). Package `survey` của R tự xử lý khi thầy khai báo đầy đủ.'),
]
for i, (k, v) in enumerate(ws):
    row = w_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FFE699')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(191, 97, 14)
    row.cells[1].text = v
d.add_paragraph()

# =========================================================
# PHẦN 6 — TÍNH TRONG R
# =========================================================
add_heading('6. Cách tính weighted prevalence trong R', level=1)

add_para('Với package `survey` của Thomas Lumley:', bold=False)
d.add_paragraph()

code_tbl = d.add_table(rows=1, cols=1)
code_tbl.style = 'Table Grid'
cd = code_tbl.rows[0].cells[0]
shade_cell(cd, 'F2F2F2')
cp = cd.paragraphs[0]
cr = cp.add_run(
'library(survey)\n\n'
'# Khai báo thiết kế khảo sát\n'
'design <- svydesign(\n'
'  ids = ~school_id,        # cluster = trường học\n'
'  strata = ~district,      # tầng = quận/huyện\n'
'  weights = ~final_weight, # trọng số đã tính\n'
'  data = df,\n'
'  nest = TRUE              # trường nằm trong quận\n'
')\n\n'
'# Tính weighted prevalence + KTC 95%\n'
'svyciprop(~anxiety, design, method = "logit", level = 0.95)\n\n'
'# Phân nhóm theo giới\n'
'svyby(~anxiety, ~gender, design, svyciprop, method = "logit")\n'
)
cr.font.name = 'Consolas'
cr.font.size = Pt(10)
d.add_paragraph()

add_para('Output sẽ cho:')
for it in [
    'Weighted prevalence ước tính (ví dụ 0,348)',
    'Standard error ĐIỀU CHỈNH cho cluster + weight',
    'KTC 95% (ví dụ [0,33; 0,37])',
    'Design effect (DEFF) nếu yêu cầu',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# PHẦN 7 — CHEN 2023 CỤ THỂ
# =========================================================
add_heading('7. Áp dụng vào Chen 2023 (Bài 7 — QT007)', level=1)

chen_tbl = d.add_table(rows=6, cols=2)
chen_tbl.style = 'Table Grid'
chens = [
    ('Thiết kế mẫu',
     'Lấy mẫu cụm: 2 quận + 1 huyện của thành phố Tự Cống (Trung Quốc) → tất cả trường THCS+THPT trong khu vực → khảo sát 63.205 HS.'),
    ('Vì sao có hiệu ứng cụm',
     'HS cùng trường có intraclass correlation (ICC) > 0. Cùng thầy cô, chương trình, văn hoá → điểm lo âu tương quan.'),
    ('Vì sao cần weighted',
     '(a) Trường khác size → xác suất mỗi HS được chọn khác; (b) Một số HS có thể không hoàn thành → non-response; (c) Muốn ngoại suy cho toàn bộ HS khu vực (không chỉ mẫu 63.205).'),
    ('Công cụ sử dụng',
     'Gói "survey" của R — điều chỉnh cả weighting và cluster effect đồng thời. Output: weighted prevalence + SE đúng + KTC 95 %.'),
    ('Kết quả báo cáo',
     'Chen 2023 ghi "Tỷ lệ có trọng số của triệu chứng trầm cảm và lo âu (N = 63.205)" — nghĩa là các tỷ lệ % trong bảng đã áp dụng weighting, có thể ngoại suy cho toàn HS TT-THPT khu vực Tự Cống.'),
    ('Ý nghĩa cho đọc giả',
     'Khi thầy đọc tỷ lệ lo âu trong bài → đây KHÔNG PHẢI là "200 HS trong 1.000 HS được khảo sát" — mà là "ước tính % HS có triệu chứng TRÊN TOÀN QUẦN THỂ mà mẫu đại diện". Con số chính xác hơn và có thể so sánh với các survey khác.'),
]
for i, (k, v) in enumerate(chens):
    row = chen_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'DEEBF7')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(31, 73, 125)
    row.cells[1].text = v
d.add_paragraph()

# =========================================================
# PHẦN 8 — KHI NÀO DÙNG / KHÔNG
# =========================================================
add_heading('8. Khi nào cần weighted prevalence?', level=1)

need_tbl = d.add_table(rows=4, cols=2)
need_tbl.style = 'Table Grid'
hdr = need_tbl.rows[0]
for i, h in enumerate(['CẦN weighted', 'KHÔNG cần']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

need = [
    ('Khảo sát quốc gia (V-NAMHS, GSHS, VNAMHS)',
     'RCT với ngẫu nhiên hoá (randomization đảm bảo weight = 1)'),
    ('Stratified cluster sampling (như Chen 2023)',
     'Nghiên cứu cắt ngang với simple random sampling'),
    ('Mẫu có over-sampling nhóm thiểu số',
     'Case-control study nhỏ, hoặc nghiên cứu định tính'),
]
for i, (w, nw) in enumerate(need):
    row = need_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run('• ' + w)
    r0.font.size = Pt(11)
    cell1 = row.cells[1]
    shade_cell(cell1, 'FCE4D6')
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run('• ' + nw)
    r1.font.size = Pt(11)
d.add_paragraph()

# =========================================================
# TÓM GỌN
# =========================================================
add_heading('9. Tóm gọn 3 ý cốt lõi', level=1)

core_tbl = d.add_table(rows=3, cols=1)
core_tbl.style = 'Table Grid'
cores = [
    ('① Weighted prevalence = tỷ lệ đã điều chỉnh bằng trọng số',
     'Công thức: p = Σ(wᵢ × yᵢ) / Σ(wᵢ). Mỗi người đóng góp khác nhau tuỳ weight (= 1 / xác suất được chọn, có thể bổ sung non-response + post-stratification).',
     'E2EFDA'),
    ('② Vì sao CẦN — khi lấy mẫu phức tạp',
     'Khi xác suất chọn không đều (trường lớn/nhỏ khác nhau) + over-sampling chủ đích + non-response khác nhóm — dùng công thức thường sẽ BỊ LỆCH. Weighting giúp ngoại suy đúng cho quần thể thực.',
     'FFF2CC'),
    ('③ Cách tính — dùng gói chuyên dụng',
     'KHÔNG tính tay. Dùng package `survey` (R), proc SURVEYFREQ/SURVEYMEANS (SAS), hoặc svy commands (Stata). Tự động xử lý cluster + weight + KTC. Chen 2023 dùng gói `survey` của R.',
     'DEEBF7'),
]
for i, (k, v, clr) in enumerate(cores):
    cell = core_tbl.rows[i].cells[0]
    shade_cell(cell, clr)
    p = cell.paragraphs[0]
    r1 = p.add_run(k + '\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)
d.add_paragraph()

# TLTK
add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Lumley T. (2004). Analysis of complex survey samples. Journal of Statistical Software, 9(8):1-19. [Bài gốc gói survey của R]',
    'Lumley T. (2010). Complex Surveys: A Guide to Analysis Using R. Wiley. [Giáo khoa chuẩn]',
    'Heeringa SG, West BT, Berglund PA. (2017). Applied Survey Data Analysis (2nd ed.). Chapman & Hall/CRC. [Giáo khoa nâng cao]',
    'Chen và cộng sự (2023). Depression and anxiety symptoms among middle and high school students in Zigong, Sichuan. BMC Psychiatry. [QT007 — Bài 7, nguồn câu hỏi]',
    'Kalton G, Flores-Cervantes I. (2003). Weighting methods. Journal of Official Statistics, 19(2):81-97. [Tổng quan các loại weight]',
    'Korn EL, Graubard BI. (1999). Analysis of Health Surveys. Wiley. [Giáo khoa cho health surveys với weighting]',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 22/04/2026 | Bối cảnh Chen 2023 (N=63.205, 2 quận + 1 huyện Tự Cống, gói survey R) verify từ tóm tắt QT007 trong thư viện.')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/Weighted_prevalence_ty_le_co_trong_so.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
