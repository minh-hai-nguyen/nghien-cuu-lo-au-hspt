# -*- coding: utf-8 -*-
"""
Thêm Phụ lục giải thích ý nghĩa số liệu vào cuối Q25 VN + EN.
Để gửi cho team SPSS/statistician hiểu rõ ý nghĩa β, R², Cohen d, fit indices...
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
VN_IN = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v3.docx"
EN_IN = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_EN_v3.docx"
VN_OUT = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v3.docx"  # in-place
EN_OUT = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_EN_v3.docx"  # in-place


def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    if color: run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_para(doc, text, size=12, bold=False, italic=False, align=None,
             indent=None, space_before=3, space_after=3):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_format(r, size=size, bold=bold, italic=italic)
    return p


def add_h(doc, text, level=1):
    if level == 1:
        return add_para(doc, text, size=13, bold=True, space_before=14, space_after=6,
                        color_=None)
    return add_para(doc, text, size=12, bold=True, italic=True, space_before=8, space_after=4)


def add_h1(doc, text):
    return add_para(doc, text, size=13, bold=True, space_before=14, space_after=6)


def add_h2(doc, text):
    return add_para(doc, text, size=12, bold=True, italic=True, space_before=8, space_after=4)


def add_table_rows(doc, headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Light Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        r = cell.paragraphs[0].add_run(h)
        set_run_format(r, size=11, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            r = cell.paragraphs[0].add_run(str(v))
            set_run_format(r, size=11)
    return t


# ============================================================
# VN APPENDIX
# ============================================================
def add_vn_appendix(doc):
    add_h1(doc, "Phụ lục — Giải thích ý nghĩa số liệu thống kê (gửi nhóm phân tích SPSS)")

    add_para(doc,
             "Phụ lục này giải thích ý nghĩa của tất cả các chỉ số thống kê được báo cáo trong bài, "
             "phục vụ trao đổi với nhóm phân tích dữ liệu (SPSS/R/Mplus). Mỗi mục đi kèm: (a) ý "
             "nghĩa khái niệm; (b) cách diễn giải; (c) ngưỡng tham chiếu; (d) ví dụ từ kết quả "
             "của bài.",
             size=12, align="justify", indent=0.5)

    # ===== 1. β chuẩn hóa =====
    add_h2(doc, "1. Hệ số β (beta) chuẩn hóa trong mô hình SEM")
    add_para(doc,
             "Ý nghĩa: β chuẩn hóa (standardized regression coefficient) biểu thị mức thay đổi "
             "trung bình của biến phụ thuộc (tính bằng độ lệch chuẩn) khi biến độc lập thay đổi "
             "một độ lệch chuẩn, sau khi các biến khác trong mô hình được kiểm soát. β nằm trong "
             "khoảng từ −1 đến +1.",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Cách diễn giải:",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["|β|", "Diễn giải", "Ý nghĩa thực tiễn"],
                   [
                       ["< 0,10", "Rất nhỏ", "Hầu như không có hiệu ứng"],
                       ["0,10 – 0,29", "Nhỏ", "Có hiệu ứng nhưng yếu"],
                       ["0,30 – 0,49", "Trung bình", "Hiệu ứng đáng kể về mặt thực tiễn"],
                       ["≥ 0,50", "Lớn", "Hiệu ứng mạnh, có ý nghĩa lâm sàng"],
                   ])
    add_para(doc,
             "Dấu (+/−): β dương cho biết quan hệ thuận (biến độc lập tăng → biến phụ thuộc tăng); "
             "β âm cho biết quan hệ nghịch (biến độc lập tăng → biến phụ thuộc giảm). Trong bài, "
             "yếu tố bảo vệ thường có β âm (tốt cho học sinh).",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Ví dụ trong bài (sắp theo cường độ):",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["Đường dẫn", "β", "Diễn giải"],
                   [
                       ["BPĐP → RLLATQ", "0,749", "RẤT LỚN dương — biện pháp đối phó kém thích nghi làm tăng mạnh lo âu tổng quát (nhưng fit kém, xem mục 4)"],
                       ["ALHT → RLLATQ", "0,510", "LỚN dương — áp lực học tập làm tăng mạnh lo âu tổng quát"],
                       ["ALHT → RLLAXH", "0,490", "LỚN dương — áp lực học tập làm tăng mạnh lo âu xã hội"],
                       ["TTr → RLLATQ", "−0,455", "LỚN âm — lòng tự trọng cao bảo vệ mạnh khỏi lo âu tổng quát"],
                       ["TTr → RLLAXH", "−0,415", "LỚN âm — lòng tự trọng bảo vệ mạnh khỏi lo âu xã hội"],
                       ["NĐT → RLLAXH", "0,383", "TRUNG BÌNH dương — nghiện điện thoại làm tăng lo âu xã hội"],
                       ["BNHĐ → RLLAC", "0,376", "TRUNG BÌNH dương — bắt nạt học đường tăng lo âu chia ly (mạnh nhất trong nhóm BNHĐ)"],
                       ["NĐT → RLLATQ", "0,336", "TRUNG BÌNH dương — nghiện điện thoại tăng lo âu tổng quát"],
                       ["HTCM → RLLAXH", "−0,273", "NHỎ-TRUNG BÌNH âm — hỗ trợ cha mẹ bảo vệ khỏi lo âu xã hội"],
                       ["ALHT → RLLAC", "0,253", "NHỎ-TRUNG BÌNH dương — áp lực học tập tăng lo âu chia ly"],
                       ["BNHĐ → RLLAXH", "0,253", "NHỎ-TRUNG BÌNH dương"],
                       ["BNHĐ → RLLATQ", "0,215", "NHỎ dương"],
                       ["GBTH → RLLA", "−0,155", "NHỎ âm — gắn bó trường học bảo vệ nhẹ"],
                       ["HTCM → RLLAC", "0,000", "BẰNG KHÔNG — hỗ trợ cha mẹ KHÔNG có tác động lên lo âu chia ly (p=0,997)"],
                   ])

    # ===== 2. R² =====
    add_h2(doc, "2. Hệ số xác định R² (R bình phương)")
    add_para(doc,
             "Ý nghĩa: R² biểu thị tỷ lệ phần trăm phương sai của biến phụ thuộc được giải thích "
             "bởi (các) biến độc lập trong mô hình. R² nằm trong khoảng từ 0 đến 1; nhân 100 để "
             "ra phần trăm.",
             size=12, align="justify", indent=0.5)
    add_para(doc, "Cách diễn giải (Cohen, 1988):", size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["R²", "Diễn giải"],
                   [
                       ["0,02 – 0,12", "Hiệu ứng nhỏ"],
                       ["0,13 – 0,25", "Hiệu ứng trung bình"],
                       ["≥ 0,26", "Hiệu ứng lớn"],
                   ])
    add_para(doc,
             "Ví dụ trong bài:",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["Mô hình", "R²", "Diễn giải"],
                   [
                       ["BPĐP → RLLATQ", "0,561", "Rất lớn — 56,1% phương sai lo âu tổng quát được giải thích bởi đối phó kém thích nghi"],
                       ["Mô hình tích hợp YTNC + YTBV → RLLA", "0,598", "Rất lớn — 59,8% phương sai lo âu được giải thích bởi tổ hợp yếu tố nguy cơ và bảo vệ"],
                       ["YTNC tổng → RLLA", "0,558", "Lớn — yếu tố nguy cơ riêng giải thích 55,8% phương sai"],
                       ["ALHT → RLLA tổng", "0,284", "Lớn — áp lực học tập một mình giải thích 28,4% phương sai"],
                       ["TTr → RLLA tổng", "0,209", "Trung bình — lòng tự trọng một mình giải thích 20,9% phương sai"],
                       ["NĐT → RLLA tổng", "0,160", "Trung bình — nghiện điện thoại giải thích 16,0% phương sai"],
                       ["YTBV tổng → RLLA", "0,124", "Nhỏ — yếu tố bảo vệ riêng giải thích 12,4% phương sai"],
                       ["BNHĐ → RLLA tổng", "0,076", "Nhỏ — bắt nạt một mình giải thích 7,6% phương sai"],
                       ["GBTH → RLLA tổng", "0,024", "Rất nhỏ"],
                   ])

    # ===== 3. Cohen d =====
    add_h2(doc, "3. Cohen d (kích thước hiệu ứng giữa hai nhóm)")
    add_para(doc,
             "Ý nghĩa: Cohen d biểu thị chênh lệch trung bình giữa hai nhóm (ví dụ nam vs nữ) "
             "tính bằng số độ lệch chuẩn của tổng thể. d không có giới hạn trên về mặt lý "
             "thuyết nhưng trong nghiên cứu hành vi thường nằm trong khoảng −2 đến +2.",
             size=12, align="justify", indent=0.5)
    add_para(doc, "Cách diễn giải (Cohen, 1988):", size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["|d|", "Diễn giải"],
                   [
                       ["< 0,20", "Rất nhỏ / không đáng kể"],
                       ["0,20 – 0,49", "Nhỏ"],
                       ["0,50 – 0,79", "Trung bình"],
                       ["≥ 0,80", "Lớn"],
                   ])
    add_para(doc,
             "Ví dụ pattern ba tầng giới trong bài:",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["Phân nhóm lo âu", "F (giới)", "p", "Cohen d", "Diễn giải"],
                   [
                       ["Lo âu tổng quát (RLLATQ)", "44,484", "<0,001", "0,365", "NHỎ — nữ > nam đáng kể"],
                       ["Lo âu xã hội (RLLAXH)", "45,984", "<0,001", "0,370", "NHỎ — nữ > nam đáng kể (gần bằng RLLATQ)"],
                       ["Lo âu chia ly (RLLAC)", "0,246", "0,620", "≈ 0,03", "KHÔNG ĐÁNG KỂ — không có chênh lệch giới"],
                   ])
    add_para(doc,
             "Phù hợp pattern toàn cầu (Salk và cộng sự, 2017): giới khuếch đại nguy cơ đồng đều "
             "cho các phân nhóm lo âu nhận thức – đánh giá (tổng quát, xã hội) nhưng không "
             "khuếch đại cho phân nhóm dựa trên gắn bó (chia ly).",
             size=12, align="justify", indent=0.5)

    # ===== 4. Fit indices =====
    add_h2(doc, "4. Chỉ số phù hợp mô hình SEM (Fit Indices)")
    add_para(doc,
             "Ý nghĩa: Các chỉ số đánh giá mức độ mà mô hình lý thuyết (giả định bởi nhóm tác "
             "giả) phù hợp với cấu trúc covariance trong dữ liệu thực. Mô hình có fit tốt là "
             "điều kiện để các hệ số β được diễn giải đáng tin cậy.",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Bốn chỉ số chính theo Hu và Bentler (1999):",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["Chỉ số", "Ngưỡng tốt", "Ngưỡng chấp nhận", "Ý nghĩa khi vượt ngưỡng kém"],
                   [
                       ["CFI", "≥ 0,95", "≥ 0,90", "< 0,90 → mô hình không phù hợp"],
                       ["TLI", "≥ 0,95", "≥ 0,90", "< 0,90 → mô hình không phù hợp"],
                       ["RMSEA", "≤ 0,06", "≤ 0,08", "> 0,10 → mô hình rất kém"],
                       ["SRMR", "≤ 0,08", "≤ 0,10", "> 0,10 → cần đánh giá lại"],
                   ])
    add_para(doc,
             "Ví dụ trong bài:",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["Mô hình", "RMSEA", "CFI", "Kết luận"],
                   [
                       ["BNHĐ → RLLA (2 nhân tố)", "0,030", "0,999", "XUẤT SẮC — được ưu tiên báo cáo"],
                       ["BNHĐ → RLLA (3 nhân tố)", "0,129", "—", "KÉM — cần thận trọng"],
                       ["BPĐP → RLLA", "0,080 – 0,204", "0,865 – 0,911", "KÉM trên mọi tổ hợp — diễn giải EXPLORATORY"],
                   ])
    add_para(doc,
             "Quan trọng: khi mô hình có fit kém (như BPĐP trong bài), hệ số β rất lớn (0,749) "
             "không thể được hiểu là hiệu ứng nhân quả một chiều. Cách diễn giải đúng là quan hệ "
             "hai chiều (bidirectional escalation): đối phó kém thích nghi vừa khuếch đại vừa "
             "được khuếch đại bởi lo âu (Compas và cộng sự, 2017).",
             size=12, align="justify", indent=0.5)

    # ===== 5. Tỷ số β =====
    add_h2(doc, "5. Tỷ số |β| (so sánh cường độ giữa các yếu tố)")
    add_para(doc,
             "Ý nghĩa: Tỷ số |β yếu tố A| / |β yếu tố B| cho biết cường độ tương đối giữa hai "
             "yếu tố cùng dự báo một kết cục. Tỷ số gần 1 → cường độ ngang nhau; tỷ số > 1 → "
             "yếu tố A mạnh hơn; tỷ số < 1 → yếu tố B mạnh hơn.",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Ví dụ trong bài (Lòng tự trọng so với Áp lực học tập):",
             size=12, bold=True, indent=0.5)
    add_table_rows(doc,
                   ["Phân nhóm lo âu", "|β TTr|", "|β ALHT|", "Tỷ số", "Diễn giải"],
                   [
                       ["Lo âu tổng quát", "0,455", "0,510", "0,892", "GẦN NGANG — tự trọng yếu hơn áp lực học tập chỉ ~10%"],
                       ["Lo âu xã hội", "0,415", "0,490", "0,847", "GẦN NGANG — tự trọng yếu hơn áp lực học tập ~15%"],
                       ["Lo âu chia ly", "0,087", "0,253", "0,344", "ÍT NGANG — tự trọng yếu hơn áp lực học tập rất nhiều"],
                   ])
    add_para(doc,
             "Phát hiện đặc thù: ở hai phân nhóm lo âu nhận thức – đánh giá (tổng quát + xã "
             "hội), lòng tự trọng có cường độ NGANG với áp lực học tập, gợi ý lòng tự trọng có "
             "thể là một mục tiêu can thiệp song song hiệu quả như giảm áp lực học tập.",
             size=12, align="justify", indent=0.5)

    # ===== 6. KTC 95% =====
    add_h2(doc, "6. Khoảng tin cậy 95% (KTC 95% / 95% CI)")
    add_para(doc,
             "Ý nghĩa: KTC 95% là khoảng giá trị mà ta có 95% tin cậy rằng giá trị tham số "
             "thật của tổng thể nằm trong đó. KTC hẹp → ước lượng chính xác; KTC rộng → ước "
             "lượng kém chính xác.",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Quy tắc kiểm định ý nghĩa: nếu KTC 95% KHÔNG chứa số 0, hệ số có ý nghĩa thống "
             "kê ở mức α = 0,05.",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Ví dụ trong bài: SMD của Cai và cộng sự (2025) = 0,17 (KTC 95% 0,06 – 0,29). "
             "Khoảng này không chứa 0 → hiệu ứng có ý nghĩa thống kê, mặc dù độ lớn nhỏ.",
             size=12, align="justify", indent=0.5)

    # ===== 7. p-value =====
    add_h2(doc, "7. Giá trị p (mức ý nghĩa thống kê)")
    add_para(doc,
             "Ý nghĩa: p-value là xác suất quan sát được kết quả hiện tại (hoặc cực đoan hơn) "
             "với giả định rằng giả thuyết không (H0) đúng. p càng nhỏ → bằng chứng càng mạnh "
             "chống H0.",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Ngưỡng thông thường: p < 0,001 (rất có ý nghĩa); p < 0,01 (có ý nghĩa cao); "
             "p < 0,05 (có ý nghĩa); p ≥ 0,05 (không có ý nghĩa thống kê).",
             size=12, align="justify", indent=0.5)
    add_para(doc,
             "Ví dụ trong bài: β HTCM → RLLAC = 0,000 với p = 0,997 → KHÔNG có ý nghĩa thống "
             "kê → hỗ trợ cha mẹ KHÔNG có tác động trực tiếp lên lo âu chia ly trong mẫu này. "
             "Đây là phát hiện đặc thù văn hóa Á.",
             size=12, align="justify", indent=0.5)

    # ===== 8. SMD, OR, r =====
    add_h2(doc, "8. SMD, OR, r — các thước đo hiệu ứng phổ biến khác")
    add_table_rows(doc,
                   ["Chỉ số", "Tên đầy đủ", "Ý nghĩa", "Ngưỡng"],
                   [
                       ["SMD", "Standardized Mean Difference", "Chênh lệch trung bình chuẩn hóa giữa hai nhóm (dùng trong phân tích tổng hợp)", "0,2 nhỏ; 0,5 trung bình; 0,8 lớn"],
                       ["OR", "Odds Ratio (tỷ số chênh)", "Xác suất xảy ra sự kiện ở nhóm phơi nhiễm chia cho nhóm không phơi nhiễm", "1 = không khác biệt; > 1 tăng nguy cơ; < 1 giảm nguy cơ"],
                       ["r", "Hệ số tương quan Pearson", "Mức độ tương quan tuyến tính giữa hai biến", "0,10 nhỏ; 0,30 trung bình; 0,50 lớn"],
                       ["g", "Hedges g", "Tương tự Cohen d nhưng hiệu chỉnh cho cỡ mẫu nhỏ", "Diễn giải như Cohen d"],
                       ["SUCRA", "Surface Under Cumulative Ranking", "Xếp hạng phương án điều trị trong phân tích mạng lưới (0–100%; cao hơn = tốt hơn)", "—"],
                   ])

    # ===== 9. Yêu cầu cho team SPSS =====
    add_h2(doc, "9. Yêu cầu cụ thể cho nhóm phân tích SPSS/lavaan/Mplus")
    add_para(doc, "Khi tái phân tích dữ liệu raw (n = 1.352), nhóm phân tích cần báo cáo:",
             size=12, bold=True, align="justify", indent=0.5)
    items = [
        "Cho mỗi mô hình SEM: χ² (df), CFI, TLI, RMSEA (kèm KTC 90%), SRMR.",
        "Cho mỗi đường dẫn: β chuẩn hóa, sai số chuẩn, KTC 95% bootstrap (5.000 lần), p-value.",
        "Cho mô hình mediation: a, b, c, c', và indirect = a × b kèm KTC 95% bootstrap.",
        "Cho multi-group invariance: configural, metric, scalar — Δχ², ΔCFI, ΔRMSEA.",
        "Cho phân tích giới: F, p, Cohen d cho mỗi phân nhóm lo âu.",
        "Cho Harman's single-factor test: % phương sai của nhân tố đầu tiên.",
        "Quy ước báo cáo: dấu phẩy thập phân kiểu Việt Nam (0,376 thay vì 0.376) cho bản tiếng Việt; dấu chấm cho bản tiếng Anh.",
        "Khi có FIT KÉM (RMSEA > 0,08 hoặc CFI < 0,90), cần phân tích so sánh mô hình thay thế (alternative model comparison).",
    ]
    for it in items:
        add_para(doc, "• " + it, size=12, indent=0.3)

    add_para(doc,
             "Phụ lục này được biên soạn nhằm chuẩn hóa cách diễn giải số liệu giữa nhóm tác "
             "giả viết bản thảo và nhóm phân tích dữ liệu, đảm bảo mọi con số trong bản thảo "
             "đều khớp với output SEM gốc và được hiểu nhất quán.",
             size=11, italic=True, align="justify", indent=0.5)


# ============================================================
# EN APPENDIX (brief)
# ============================================================
def add_en_appendix(doc):
    add_h1(doc, "Appendix — Statistical glossary for data analysts")

    add_para(doc,
             "This appendix decodes every statistical index reported in the manuscript for "
             "communication with the data-analysis team (SPSS/lavaan/Mplus). Each entry covers "
             "(a) conceptual meaning, (b) interpretation, (c) reference thresholds, and (d) an "
             "example from the present study.",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "1. Standardized beta (β) in SEM")
    add_para(doc,
             "Meaning: Standardized regression coefficient indicating the change in the dependent "
             "variable (in standard-deviation units) per one-SD increase in the predictor, "
             "controlling for other predictors. β ranges from −1 to +1.",
             size=12, align="justify", indent=0.5)
    add_table_rows(doc,
                   ["|β|", "Interpretation"],
                   [
                       ["< 0.10", "Negligible"],
                       ["0.10 – 0.29", "Small"],
                       ["0.30 – 0.49", "Medium"],
                       ["≥ 0.50", "Large"],
                   ])
    add_para(doc,
             "Examples (sorted by magnitude): β BPĐP → GAD = 0.749 (very large positive, but "
             "poor fit — see §4); β academic pressure → GAD = 0.510 (large positive); β "
             "self-esteem → GAD = −0.455 (large negative — protective); β school bullying → "
             "separation anxiety = 0.376 (medium positive, strongest in bullying model); β "
             "parental support → separation anxiety = 0.000 (null direct effect, p = 0.997).",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "2. Coefficient of determination R²")
    add_para(doc,
             "Meaning: Proportion of variance in the dependent variable explained by the "
             "predictor(s). Cohen (1988): R² ≥ 0.26 = large effect. Examples: integrated risk + "
             "protective model R² = 0.598 (very large); academic pressure alone R² = 0.284 "
             "(large); self-esteem alone R² = 0.209 (medium); bullying alone R² = 0.076 (small).",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "3. Cohen d (between-group effect size)")
    add_para(doc,
             "Meaning: Mean-difference between two groups in pooled SD units. Cohen (1988): "
             "0.2 small, 0.5 medium, 0.8 large. Three-tier gender pattern in the present study: "
             "GAD d = 0.365 (small significant), social anxiety d = 0.370 (small significant), "
             "separation anxiety d ≈ 0.03 (negligible, non-significant). Consistent with Salk "
             "et al. (2017) global pattern.",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "4. SEM fit indices (Hu and Bentler, 1999)")
    add_table_rows(doc,
                   ["Index", "Good fit", "Acceptable", "Poor fit interpretation"],
                   [
                       ["CFI", "≥ 0.95", "≥ 0.90", "< 0.90 → model not supported"],
                       ["TLI", "≥ 0.95", "≥ 0.90", "< 0.90 → model not supported"],
                       ["RMSEA", "≤ 0.06", "≤ 0.08", "> 0.10 → very poor"],
                       ["SRMR", "≤ 0.08", "≤ 0.10", "> 0.10 → reassess"],
                   ])
    add_para(doc,
             "Example: Bullying → anxiety (two-factor) showed excellent fit (RMSEA = 0.030; "
             "CFI = 0.999); the three-factor disaggregated model showed marginal fit (RMSEA = "
             "0.129). The maladaptive coping → anxiety models showed poor fit (RMSEA 0.080–"
             "0.204; CFI 0.865–0.911) across all configurations — therefore the β = 0.749 "
             "coefficient is interpreted as exploratory and most plausibly reflects "
             "bidirectional escalation rather than unidirectional causation.",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "5. Magnitude ratio |β| (cross-factor comparison)")
    add_para(doc,
             "Computed as |β factor A| / |β factor B| to compare the strength of two predictors "
             "of the same outcome. In the present study, |β self-esteem| / |β academic pressure| "
             "yields 0.892 for GAD, 0.847 for social anxiety, and 0.344 for separation anxiety, "
             "indicating that self-esteem rivals academic pressure in the two cognitive-"
             "evaluative subtypes but not in the attachment-based subtype.",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "6. Requirements for the data-analysis team")
    items = [
        "For each SEM: χ² (df), CFI, TLI, RMSEA (with 90% CI), SRMR.",
        "For each pathway: standardized β, SE, 95% bootstrap CI (5,000 resamples), p-value.",
        "For mediation: a, b, c, c′ coefficients and indirect = a × b with 95% bootstrap CI.",
        "For multi-group invariance: configural, metric, scalar — Δχ², ΔCFI, ΔRMSEA.",
        "For gender analyses: F, p, Cohen d for each anxiety subtype.",
        "For common-method variance: Harman's single-factor test (% variance of first factor).",
        "Where fit is poor (RMSEA > 0.08 or CFI < 0.90), report alternative-model comparison.",
        "Use period for decimal separators in the English manuscript; commas for the Vietnamese manuscript.",
    ]
    for it in items:
        add_para(doc, "• " + it, size=12, indent=0.3)


# ============================================================
# RUN
# ============================================================
def main():
    # VN
    print("Adding VN appendix...")
    doc = Document(VN_IN)
    add_vn_appendix(doc)
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""
    cp.subject = ""; cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 5, 1, 9, 0, 0)
    cp.modified = datetime(2026, 5, 15, 8, 0, 0)
    doc.save(VN_OUT)
    print(f"  Saved: {VN_OUT}")

    # EN
    print("Adding EN appendix...")
    doc = Document(EN_IN)
    add_en_appendix(doc)
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""
    cp.subject = ""; cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 5, 1, 9, 0, 0)
    cp.modified = datetime(2026, 5, 15, 8, 0, 0)
    doc.save(EN_OUT)
    print(f"  Saved: {EN_OUT}")
    print("\n[DONE]")


if __name__ == "__main__":
    main()
