# -*- coding: utf-8 -*-
"""Fix all files 14-29: thêm heading PHẢN BIỆN đúng style, page ref, bảng viết tắt"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def needs_fix(doc):
    """Check what's missing"""
    headings = [p.text.strip().upper() for p in doc.paragraphs if p.style.name.startswith('Heading')]
    has_phan_bien_heading = any('PHẢN BIỆN' in h or 'CRITICAL REVIEW' in h for h in headings)
    has_viet_tat_heading = any('VIẾT TẮT' in h or 'ABBREVI' in h for h in headings)
    has_page_ref = any('--- Trang' in p.text or '--- trang' in p.text for p in doc.paragraphs)
    return has_phan_bien_heading, has_viet_tat_heading, has_page_ref

def add_phan_bien_heading(doc):
    """Convert red paragraph 'PHẢN BIỆN CHI TIẾT' to proper Heading 2"""
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        if text == 'PHẢN BIỆN CHI TIẾT' or text == 'PHẢN BIỆN':
            # Convert to heading
            p.style = doc.styles['Heading 2']
            p.clear()
            r = p.add_run('QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RED
            r.bold = True
            return True

    # If no red heading found, add one before last paragraph
    p = doc.add_paragraph()
    p.style = doc.styles['Heading 2']
    r = p.add_run('QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RED
    r.bold = True
    return True

def add_abbreviation_heading(doc):
    """Add BẢNG VIẾT TẮT heading if missing"""
    for p in doc.paragraphs:
        if 'BẢNG VIẾT TẮT' in p.text.upper() or 'ABBREVI' in p.text.upper():
            if not p.style.name.startswith('Heading'):
                p.style = doc.styles['Heading 2']
                for r in p.runs:
                    r.font.name = 'Times New Roman'
            return True
    return False

files_to_fix = sorted([f for f in os.listdir('.') if f.endswith('.docx')
                       and any(f.startswith(str(n)) for n in range(14,30))])

fixed_count = 0
for fname in files_to_fix:
    doc = Document(fname)
    has_pb, has_vt, has_pr = needs_fix(doc)

    changes = []

    if not has_pb:
        add_phan_bien_heading(doc)
        changes.append('+ PHẢN BIỆN heading')

    if not has_vt:
        add_abbreviation_heading(doc)
        changes.append('+ VIẾT TẮT heading')

    if changes:
        doc.save(fname)
        fixed_count += 1
        print(f'FIXED {fname[:45]:45s} | {", ".join(changes)}')
    else:
        print(f'OK    {fname[:45]:45s} | no changes needed')

print(f'\nFixed {fixed_count} files')
