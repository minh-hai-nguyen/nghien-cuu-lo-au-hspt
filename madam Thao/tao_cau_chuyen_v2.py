# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(8)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0,0,0)

def p(text, italic=False, size=13):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.italic = italic

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_col_width(cell, width_cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(width_cm * 567)))
    w.set(qn('w:type'), 'dxa')
    tcW.append(w)

# ===== TITLE =====
title = doc.add_heading('', level=0)
r = title.add_run(
    'HÀNH TRÌNH ĐÓNG GÓP\n'
    'TS. Nguyễn Thị Phương Thảo\n'
    'Chủ tịch HĐQT Vietjet'
)
r.font.name = 'Times New Roman'
r.font.size = Pt(16)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p('Trích từ cuộc họp xây dựng chương trình đào tạo lãnh đạo — Hội nhập & Tỏa sáng.\n'
  'Lần đầu tiên Chủ tịch kể lại hành trình đóng góp xuyên suốt mấy chục năm.', italic=True)
doc.add_paragraph()

# ===== THÔNG ĐIỆP =====
h('Thông điệp mở đầu')
p('"Mỗi con người của chúng ta đều có thể trở nên vĩ đại. '
  'Làm những cái việc lớn có ý nghĩa, mà từ những cái hoạt động rất là đơn giản, giản dị của mình thôi."', italic=True)
p('"Tất cả những cái hoạt động của mình cho người dân, nó rất là thiết thực, ý nghĩa. '
  'Nó xuyên suốt mấy chục năm. Bây giờ mình phải lan tỏa tinh thần ấy cho nhân viên của mình, '
  'cán bộ của mình, để mọi người tự hào, mọi người tự tin rằng mình sẽ làm được, '
  'bất chấp tất cả những khó khăn, thách thức từ bên ngoài."')
doc.add_paragraph()

# ===== PHẦN I: HÀNH TRÌNH KINH TẾ =====
h('Phần I — Từ xóa nợ đến mở cửa kinh tế')

p('Chủ tịch kể lại hành trình từ thời kỳ Việt Nam còn là "đất nước con nợ, nợ xấu":')
p('"Vận động Nga xóa 95% nợ — mấy tỷ đô la. Còn 5% thì trả bằng hàng hóa. '
  'Việt Nam đi trả bằng những đồ thủ công mỹ nghệ, gạo, cao su. '
  'Thế là Việt Nam mình mới thoát ra khỏi cái đất nước là con nợ, nợ xấu ấy, bắt đầu đi ra thế giới."')
p('"Rồi mình lại đi tư vấn phát hành những cái trái phiếu chính phủ đầu tiên '
  'ở thị trường chứng khoán New York. Thành công những cái tỷ đô la đầu tiên. '
  'Rồi thành lập những cái ngân hàng thương mại đầu tiên, '
  'bắt đầu lúc mà ngành ngân hàng mới hình thành và phát triển."')
p('"Kinh tế bắt đầu mở ra cái đổi mới đấy, kinh tế thị trường. '
  'Trong cái thời gian này đất nước thì thiếu thốn — '
  'mình cung cấp bao nhiêu là các loại thiết bị, phân bón, hàng hóa, nhu yếu phẩm. '
  'Rồi hỗ trợ cho Việt Nam xuất khẩu các nông sản ra nước ngoài."')
p('"Mình đi theo tất cả những cái chương trình lớn quốc tế, '
  'kiểu như chương trình quốc tế của Liên Hiệp Quốc, '
  'và các chương trình tài chính của các quốc gia. '
  'Đến lúc mình có cái dự án, có cái luật về hàng không — thì là gặp cái nhân duyên đầu tiên."')
doc.add_paragraph()

# ===== PHẦN II: GIẢI CỨU CHỨNG KHOÁN =====
h('Phần II — Giải cứu sàn chứng khoán')

p('Khi sàn chứng khoán Việt Nam bị treo hệ thống, Chủ tịch đã tham gia giải cứu:')
p('"Giải cứu cho cả cái thị trường vốn của Việt Nam chứ. '
  'Cả nước Mỹ có sàn New York, Nasdaq — 2 cái sàn chứng khoán. '
  'Cả nước Anh có London Stock Exchange. '
  'Mà nó bị treo thì giải cứu mấy trăm tỷ đô la ở đấy."')
p('"Có hơn 100 công ty chứng khoán hoạt động. '
  '4 triệu khách hàng là các tổ chức và cá nhân trong nước và ngoài nước. '
  'Mỗi ngày giao dịch khoảng 1 tỷ đô la, mà cao là 2 tỷ đô la. '
  'Mình cũng nhảy vào giải cứu."')
p('"Bao nhiêu người được hưởng lợi từ đó. '
  'Chính phủ thì được cái sàn, được cả nước. '
  'Các công ty, các nhà đầu tư thì được hưởng lợi. '
  'Bản thân các công ty niêm yết ở đấy thì người ta có thị trường vốn để huy động vốn, hoạt động, phát triển. '
  'Thế còn mình thì không cần cái gì cả."')
p('Một thành viên nhóm bổ sung: giải cứu chứng khoán liên quan đến FPT — '
  '"bây giờ nó vẫn được thu tiền từ cái đó của mình ấy."')
doc.add_paragraph()

# ===== PHẦN III: HÀNG KHÔNG =====
h('Phần III — Hàng không và sân bay Tân Sơn Nhất')

p('"Năm 2013 mình làm, 2014 bắt đầu chạy. '
  'Dự án nâng cao năng lực điều hành cất hạ cánh ở sân bay Tân Sơn Nhất — '
  'nâng lên gấp đôi từ 30 lượt đến 60 lượt. Bây giờ đang vận hành 48 lượt. '
  'Từ 24 triệu lượt khách lên 42 triệu lượt khách."')
p('"Dự án ấy làm cho cả ngành hàng không. '
  'ACV mười mấy năm liền, mỗi năm đón 16 triệu lượt khách, '
  'lợi nhuận mấy trăm triệu đô la một năm. Chứ không phải là doanh số đâu. '
  'Nghĩa là họ đóng góp cho ngân sách. '
  'Thì làm sao mà chẳng có tiền đi xây sân bay."')
doc.add_paragraph()

# ===== PHẦN IV: NGOẠI GIAO =====
h('Phần IV — Ngoại giao qua mua máy bay')

p('"Mua máy bay Boeing là đi với nước Mỹ — ông Trump, ông Obama, ông Bush. '
  'Mua máy bay Airbus là đi với Pháp, Đức, Tây Ban Nha — những nước sản xuất Airbus. '
  'Mua máy bay COMAC của Trung Quốc là chương trình của ông Tập Cận Bình."')
p('"Là mình tạo nên xây dựng mối quan hệ ngoại giao, xây dựng kinh tế giữa các quốc gia, dân tộc. '
  'Mình sẽ giảm bớt đi những chiến tranh, những sự xung đột."')
p('"Máy bay mình bay đến đâu, biên giới đất nước nó mở ra. '
  'Mình mang cờ cho khắp thế giới, các đường bay mình đến."')
doc.add_paragraph()

# ===== PHẦN V: APEC =====
h('Phần V — APEC và năng lực triển khai thần tốc')

p('Vietjet phục vụ APEC 2 lần, năm 2017 — "Đưa Việt Nam vào thế giới và đưa thế giới về Việt Nam."')
p('"Cái giải pháp chiếu sáng mình làm có 90 ngày thôi. '
  'Bình thường nó phải mất 12 tháng mới xong. '
  'Trong khi đó Hàn Quốc triển khai 12 năm mà không chạy, '
  'đến lúc bàn giao đã 17 năm. Tốn 100 triệu đô rồi, không bàn giao được. '
  'Tại vì họ phải chạy thôi chứ không là nó lại thành giống như Long Thành."')
p('"Như là cái tòa Marina này, nó đã mang đến một cái sinh khí cho cả cái thành phố này. '
  'Xây cung hội nghị 4 tháng."')
p('"Ngày 9 tháng 1 năm ngoái, không ai mời mình đâu, nhưng mà thấy thời cuộc nó căng thẳng. '
  'Mình có mối quan hệ tốt với ông Trump, Đại sứ quán Mỹ mời mình, '
  'thì mình bay luôn, mang máy bay đi luôn. Cũng tạo nên thiện cảm, cũng đỡ căng thẳng."')
p('Chủ tịch cũng nhắc đến đóng góp liên quan đến UNESCO và cột cờ Hoàng Thành Thăng Long — '
  '"Cái ý nghĩa xã hội mình mang đến là cái cột cờ Hà Nội ấy."')
doc.add_paragraph()

# ===== PHẦN VI: XÃ HỘI =====
h('Phần VI — Hoạt động xã hội và thiện nguyện')

p('Chủ tịch kể lại hàng loạt hoạt động xã hội xuyên suốt mấy chục năm:')
p('"Cờ vua Việt Nam lên top 7 thế giới. '
  'Làm bóng đá futsal. Rồi bóng đá nữ, rồi thể thao người khuyết tật."')
p('"Các chương trình mổ mắt, mổ tim, khám bệnh, phát thuốc. '
  'Phát bảo hiểm y tế. Phát cứu trợ."')
p('Một thành viên nhóm bổ sung chi tiết: "Riêng cái chương trình mổ mắt cho người nghèo — '
  'một năm 10.000 ca mổ mắt. Rồi tổ chức khám chữa bệnh ở khắp nơi."')
p('"Hay là những cái phong trào mà cũng chẳng có ai kêu gọi — '
  'kêu gọi cho bà con bắt đầu đưa hàng hóa lên trên TikTok để bán ấy. '
  'Thực ra mình là những cái người đi phát động."')
doc.add_paragraph()

# ===== PHẦN VII: COVID =====
h('Phần VII — Mùa COVID')

p('"Mùa Covid mình lao đi chở vắc-xin, thiết bị y tế, bao nhiêu những bao gạo. '
  'Không có những cái đó thì làm sao mà người ta được yên lòng."')
p('Một thành viên bổ sung: "Xe cứu thương, máy thở, nấu ăn, tiêm chủng, xét nghiệm, vắc-xin các loại."')
doc.add_paragraph()

# ===== PHẦN VIII: THIẾT KẾ =====
h('Phần VIII — Triết lý thiết kế vì con người')

p('"Tất cả những cái gì mà mình làm — thiết kế một cái nhà chung cư, '
  'là phải có cái thang máy lớn để dành cho cáng cứu thương và quan tài. '
  'Từ cái thiết kế là mình cũng nghĩ đến con người — '
  'lúc người ta bệnh, người ta cần phải đưa đi cáng. '
  'Một cái thang máy ở nhà chung cư mà không thiết kế ra như thế."')
p('"Ở đâu có bậc thang là mình phải có cái đường dốc cho xe nôi, cho xe người khuyết tật. '
  'Đấy là tinh thần mình đưa vào mọi hoạt động."')
doc.add_paragraph()

# ===== PHẦN IX: TINH THẦN VÀ TRIẾT LÝ =====
h('Phần IX — Tinh thần lãnh đạo')

h('Cái hồn của đào tạo', 2)
p('"Đấy là cái tinh thần mà mình đưa vào cho cái đào tạo của mình. '
  'Nó khác với những cái hoạt động thông thường mà sau này mình cứ làm cho nó mai một đi, '
  'làm cho nó đi đâu mất. '
  'Tức là nó thiếu cái kiểu giống như con người, nó thiếu cái hồn của nó. '
  'Chính những cái tinh thần đấy nó mới là tạo cho người ta cái động lực."')

h('Hình ảnh tổ bay', 2)
p('"Ngày xưa tổ bay là briefing ở tại công ty, thậm chí là còn xếp hàng ngay ngắn. '
  'Đi một lượt, kiểm tra lại với nhau, xong rồi diễu ở trong văn phòng ấy. '
  'Nhân viên đầu tiên nhìn cái tổ bay này chuẩn bị làm nhiệm vụ, '
  'thì người ta tự hào, người ta thấy cái giá trị, cái ý nghĩa của công việc."')
p('"Chứ còn bây giờ là cứ đầu tắt mặt tối, ra khỏi nhà 2-3 giờ sáng, '
  'ra sân bay rồi về đi ngủ. Kiếm một số tiền thì không nhỏ, nhưng mà cũng đủ sống. '
  'Không có thời gian tiêu."')
p('"Người ta rời đi không phải vì chế độ của mình. Chế độ của mình là chế độ tốt. '
  'Nhưng mà cái môi trường, người ta cảm thấy là người ta không có ý nghĩa."')

h('Nhân quả', 2)
p('"Mình tin vào nhân quả. Nhân nào quả đấy. Mình làm chưa tốt thì cuộc sống sẽ có cái quả tương ứng. '
  'Chứ không cần phải có hành động gì xử lý căng thẳng."')
p('"Cuộc sống, người ta là con người, người ta có quyền sai. '
  'Thì mình hãy tạo cho người ta cơ hội để có thể khắc phục."')

h('Công bằng', 2)
p('"Tạo nên sự công bằng trong tổ chức là cần thời gian, cần cơ chế. '
  'Chứ không phải bây giờ cấp dưới đụng chạm một cái gì đấy mà mình cho rằng là sai, '
  'thì mình xử lý rất là mạnh. '
  'Trong khi đó lãnh đạo ở trên không gương mẫu, thì đấy nó không công bằng."')

h('Coaching', 2)
p('"Một trong những tinh thần mình mang đến đây với các lãnh đạo là tinh thần học thật. '
  'Cái tinh thần phải trân trọng đào tạo. '
  'Mà lúc nào đó nó mai một đi thì quay trở lại. '
  'Và mỗi người phải trở thành người coach cho những người xung quanh. '
  'Cũng phải biết cách học hỏi."')
doc.add_paragraph()

# ===== KẾT =====
h('Thông điệp kết thúc')
p('"Mỗi con người đều có thể trở nên vĩ đại. '
  'Làm những việc lớn có ý nghĩa, mà từ những hoạt động rất là đơn giản, giản dị của mình thôi."', italic=True)
p('Một thành viên nhóm đề xuất: "Madame cứ nói đi, bây giờ mình quay phim. '
  'Các bạn lo bản chân, bản mang, phim và câu chuyện."')

doc.save('Câu chuyện Madame Thảo.docx')
print('Done')
