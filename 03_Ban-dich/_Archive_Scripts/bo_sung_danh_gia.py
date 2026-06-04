# -*- coding: utf-8 -*-
"""Bổ sung TÓM TẮT ĐÁNH GIÁ NHANH cho 5 bản dịch + phản biện đỏ cho V-NAMHS"""
import sys, io
try:
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except:
    pass
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def make_para(text, bold=False, color_hex=None, size=12):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman')
    rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(size * 2))
    rPr.append(sz)
    if bold:
        rPr.append(OxmlElement('w:b'))
    if color_hex:
        c = OxmlElement('w:color')
        c.set(qn('w:val'), color_hex)
        rPr.append(c)
    r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set(qn('xml:space'), 'preserve')
    r.append(t)
    p.append(r)
    return p

def insert_eval(doc, anchor_text, highlights, limitations, improvements):
    for i, para in enumerate(doc.paragraphs):
        if anchor_text in para.text:
            parent = para._element.getparent()
            idx = list(parent).index(para._element)
            elems = [make_para('')]
            elems.append(make_para('TOM TAT DANH GIA NHANH', bold=True, color_hex='0070C0', size=14))
            elems.append(make_para(''))
            elems.append(make_para('Diem noi bat:', bold=True, color_hex='0070C0'))
            for h in highlights:
                elems.append(make_para('  ' + h, color_hex='0070C0', size=11))
            elems.append(make_para(''))
            elems.append(make_para('Han che:', bold=True, color_hex='0070C0'))
            for l in limitations:
                elems.append(make_para('  ' + l, color_hex='0070C0', size=11))
            elems.append(make_para(''))
            elems.append(make_para('Huong cai thien:', bold=True, color_hex='0070C0'))
            for imp in improvements:
                elems.append(make_para('  ' + imp, color_hex='0070C0', size=11))
            elems.append(make_para(''))
            for j, elem in enumerate(elems):
                parent.insert(idx + 1 + j, elem)
            return True
    return False

def append_red(doc, lines):
    body = doc.element.body
    for text, bold, size in lines:
        body.append(make_para(text, bold=bold, color_hex='FF0000', size=size))

# === 1. HOA 2024 ===
try:
    doc = Document('02_Hoa_2024_Frontiers.docx')
    insert_eval(doc, 'COVID-19',
        ['Co mau lon (n=3.910), phuong phap hon hop GAD-7 + phong van sau 20 HS.',
         'Ty le lo au 40,6% — cao hon nhieu nuoc cung thoi diem.',
         'Phat hien dinh tinh bo sung: ap luc thi cu, ky thi xa hoi, ky vong gia dinh.',
         'Do tin cay thang do rat cao (Cronbach alpha = 0,916).'],
        ['Chi tai Ha Noi (do thi), khong dai dien nong thon va mien nui.',
         'Thoi diem khao sat trung COVID-19 lockdown (10-11/2021).',
         'GAD-7 la sang loc, khong phai chan doan — V-NAMHS chi ghi nhan 2,3% theo DISC-5.',
         'Phong van dinh tinh chi 20 HS — chua du tong quat hoa.'],
        ['Mo rong ra nong thon, mien nui, dan toc thieu so.',
         'So sanh sang loc (GAD-7) voi chan doan (DISC-5) tren cung mau.',
         'Nghien cuu doc theo doi thay doi lo au truoc-trong-sau COVID.',
         'Danh gia vai tro trung gian cua giac ngu va mang xa hoi.']
    )
    doc.save('02_Hoa_2024_Frontiers.docx')
    sys.stderr.write('Hoa 2024: OK\n')
except Exception as e:
    sys.stderr.write(f'Hoa 2024: ERROR {e}\n')

# === 2. GBD ASEAN ===
try:
    doc = Document('03_GBD_ASEAN_2025_Lancet.docx')
    insert_eval(doc, 'Gates',
        ['Lancet Public Health (IF ~72) — tap chi hang dau the gioi.',
         'Du lieu GBD chuan hoa 30 nam (1990-2021), 10 quoc gia ASEAN.',
         'Nhom 10-14 tuoi chiu ganh nang cao nhat: 16,3% tong DALYs.',
         '80,4 trieu ca roi loan tam than, tang 70% so voi 1990.'],
        ['Du lieu GBD phu thuoc chat luong bao cao quoc gia.',
         'Khong phan tich rieng hoc sinh THCS/THPT.',
         'Ty le thap o Viet Nam (10,1%) co the phan anh thieu du lieu.',
         'Mo hinh DisMod-MR dua tren gia dinh phan bo.'],
        ['Can nghien cuu cap quoc gia voi phuong phap chuan hoa tai tung nuoc ASEAN.',
         'Phan tich rieng nhom hoc sinh THCS/THPT thay vi chi theo nhom tuoi.',
         'Bo sung du lieu tu Myanmar, Lao, Campuchia.',
         'So sanh ty le GBD voi ket qua khao sat truc tiep (V-NAMHS, GSHS).']
    )
    doc.save('03_GBD_ASEAN_2025_Lancet.docx')
    sys.stderr.write('GBD ASEAN: OK\n')
except Exception as e:
    sys.stderr.write(f'GBD ASEAN: ERROR {e}\n')

# === 3. ZHAMEDEN ===
try:
    doc = Document('04_Zhameden_2025_PLOSONE.docx')
    insert_eval(doc, 'contextual',
        ['Tong quan he thong dau tien tap trung vao can thiep truong hoc tai LMIC.',
         'Phat hien: can thiep hieu qua cho tram cam (3/4) nhung kem cho lo au (1/4).',
         'Danh gia GRADE cho thay chat luong bang chung rat thap.',
         'Xac dinh khoang trong lon: chua co RCT nao tu nuoc thu nhap thap.'],
        ['Chi 6 RCTs dap ung tieu chi — so luong rat it.',
         'Khong co nghien cuu nao tu nuoc thu nhap thap.',
         'Thieu theo doi dai han (>3 thang).',
         '4/6 dung CBT — thieu da dang phuong phap.'],
        ['RCT can thiep tai truong hoc Viet Nam.',
         'Thu nghiem can thiep phi CBT: mindfulness, SEL, giao duc tam ly.',
         'Theo doi hieu qua dai han (6-12 thang).',
         'Danh gia chi phi-hieu qua cho boi canh LMIC.']
    )
    doc.save('04_Zhameden_2025_PLOSONE.docx')
    sys.stderr.write('Zhameden: OK\n')
except Exception as e:
    sys.stderr.write(f'Zhameden: ERROR {e}\n')

# === 4. ANDERSON ===
try:
    doc = Document('05_Anderson_2025_Wiley.docx')
    insert_eval(doc, 'trajectories',
        ['Tong quan 61 bai bao — bao quat nhieu yeu to: hoc tap, mang xa hoi, gia dinh.',
         '31,9% VTN 13-18 tuoi co roi loan lo au, The he Z cao nhat 3 the he.',
         '48/52 nghien cuu xac nhan tuong quan duong giua ap luc hoc tap va SKTT kem.',
         'De xuat can thiep da tang: lam sang + giao duc + cong dong.'],
        ['Tong quan tuong thuat (khong phai he thong) — khong co phan tich tong hop.',
         'Du lieu chu yeu tu phuong Tay — thieu boi canh chau A, chau Phi.',
         'Khong danh gia chat luong tung nghien cuu duoc dua vao.',
         'Thieu phan tich theo nhom tuoi cu the (THCS vs THPT).'],
        ['Tong quan he thong + meta-analysis voi tieu chi PRISMA.',
         'Bo sung nghien cuu tu chau A (Viet Nam, Trung Quoc, Han Quoc).',
         'Phan tich yeu to dac thu van hoa: hoc them, ky vong gia dinh kieu A Dong.',
         'Danh gia hieu qua can thiep theo tung yeu to nguy co.']
    )
    doc.save('05_Anderson_2025_Wiley.docx')
    sys.stderr.write('Anderson: OK\n')
except Exception as e:
    sys.stderr.write(f'Anderson: ERROR {e}\n')

# === 5. V-NAMHS (eval + red critique) ===
try:
    doc = Document('06_VNAMHS_2022.docx')
    insert_eval(doc, 'policy',
        ['Khao sat dai dien quoc gia DAU TIEN tai Viet Nam dung cong cu chan doan DISC-5/DSM-5.',
         '5.996 VTN 10-17 tuoi tai 38 tinh — pham vi bao phu chua tung co.',
         'Phan biet ro van de SKTT (21,7%) va roi loan tam than (3,3%).',
         'Phat hien khoang cach dich vu lon: chi 8,4% co van de SKTT tiep can dich vu.'],
        ['Du lieu thu thap 2021-2022, co the chiu anh huong COVID-19.',
         'DISC-5 chua duoc chuan hoa cho van hoa Viet Nam.',
         'Thieu phan tich theo vung dia ly chi tiet.',
         'Chi 6 roi loan duoc danh gia — thieu OCD, roi loan an uong.'],
        ['V-NAMHS lan 2 (nghien cuu doc) de theo doi xu huong sau COVID.',
         'Chuan hoa DISC-5 cho boi canh van hoa Viet Nam.',
         'Phan tich chi tiet theo vung: do thi, nong thon, mien nui, DTTS.',
         'RCT can thiep tai truong hoc dua tren ket qua V-NAMHS.']
    )
    # Add red critique at end
    append_red(doc, [
        ('', False, 12),
        ('PHAN BIEN CHI TIET', True, 14),
        ('', False, 12),
        ('Diem manh:', True, 12),
        ('  Khao sat quoc gia dau tien voi DISC-5/DSM-5. n=5.996, 38 tinh.', False, 11),
        ('  Phan biet van de SKTT (21,7%) va roi loan (3,3%) — giup hoach dinh chinh sach.', False, 11),
        ('  Danh gia su dung dich vu (8,4%) — cung cap bang chung khoang trong.', False, 11),
        ('', False, 12),
        ('Han che:', True, 12),
        ('  DISC-5 chua chuan hoa cho van hoa VN — bieu hien trieu chung co the khac phuong Tay.', False, 11),
        ('  Thu thap du lieu 2021-2022 trung COVID-19.', False, 11),
        ('  Chi 6/hang tram roi loan tam than duoc danh gia.', False, 11),
        ('  Phu huynh la gatekeeper — chi 5,1% nhan ra con can giup do.', False, 11),
        ('', False, 12),
        ('Khoang trong nghien cuu (Gap):', True, 12),
        ('  V-NAMHS lan 2 (nghien cuu doc) de theo doi xu huong sau COVID-19.', False, 11),
        ('  Chuan hoa DISC-5 cho boi canh van hoa Viet Nam.', False, 11),
        ('  Phan tich theo vung: do thi vs nong thon vs mien nui vs DTTS.', False, 11),
        ('  RCT can thiep tai truong hoc dua tren ket qua V-NAMHS.', False, 11),
        ('  So sanh ty le V-NAMHS voi GSHS va DASS-21 tren cung mau.', False, 11),
    ])
    doc.save('06_VNAMHS_2022.docx')
    sys.stderr.write('V-NAMHS: OK\n')
except Exception as e:
    sys.stderr.write(f'V-NAMHS: ERROR {e}\n')
