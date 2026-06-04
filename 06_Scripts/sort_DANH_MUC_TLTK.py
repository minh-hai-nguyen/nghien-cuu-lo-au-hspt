"""Sap xep file 01_Gui H_DANH MUC TAI LIEU THAM KHAO.docx
- Tach phan Tieng Viet va Tieng Anh
- Sort alphabet theo tac gia trong tung phan
- GIU NGUYEN mau chu, font, italic, bold, ... cua moi entry
- Output: file moi voi suffix "_DA_SAP_XEP"
"""
import sys, io, re, unicodedata
from copy import deepcopy
from pathlib import Path
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

SRC = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/01_Gửi H_DANH MỤC TÀI LIỆU THAM KHẢO.docx')
OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/01_Gửi H_DANH MỤC TÀI LIỆU THAM KHẢO_DA_SAP_XEP.docx')

def remove_diacritics(s):
    """Strip Vietnamese diacritics for sort key."""
    nfkd = unicodedata.normalize('NFKD', s)
    out = ''.join(c for c in nfkd if not unicodedata.combining(c))
    # Handle đ/Đ which NFKD doesn't split
    out = out.replace('đ', 'd').replace('Đ', 'D')
    return out

def sort_key(text):
    """Build a sort key from paragraph text — strip leading punctuation/whitespace,
    lowercase + remove diacritics."""
    t = text.strip()
    # Strip leading bullet chars or numbering
    t = re.sub(r'^[\s•\-\d\.\)\(]+', '', t)
    t = remove_diacritics(t).lower().strip()
    return t

def main():
    doc = Document(SRC)
    paras = doc.paragraphs

    # Find boundaries
    idx_tlth = None  # "TÀI LIỆU THAM KHẢO"
    idx_vn = None     # "Tiếng Việt"
    idx_en = None     # "Tiếng Anh"
    for i, p in enumerate(paras):
        t = p.text.strip()
        if t == 'TÀI LIỆU THAM KHẢO': idx_tlth = i
        elif t == 'Tiếng Việt': idx_vn = i
        elif t == 'Tiếng Anh': idx_en = i

    print(f'Boundaries: TLTH={idx_tlth}, VN={idx_vn}, EN={idx_en}')
    if idx_vn is None or idx_en is None:
        print('ERROR: cannot find boundaries')
        return

    # Collect VN entries (between idx_vn+1 and idx_en-1)
    # Collect EN entries (between idx_en+1 and end)
    # Skip empty paragraphs but keep them at the end
    def collect_entries(start, end):
        entries = []  # list of (sort_key, paragraph_element_clone)
        for i in range(start, end):
            p = paras[i]
            t = p.text.strip()
            if not t:
                continue  # skip empty
            entries.append((sort_key(t), deepcopy(p._element)))
        return entries

    vn_entries = collect_entries(idx_vn + 1, idx_en)
    en_entries = collect_entries(idx_en + 1, len(paras))

    print(f'Collected: VN={len(vn_entries)} entries, EN={len(en_entries)} entries')

    # Sort
    vn_entries.sort(key=lambda x: x[0])
    en_entries.sort(key=lambda x: x[0])

    # Build new document by REPLACING body content
    # Strategy: keep paragraphs [0..idx_en], remove rest, then append sorted EN
    # Then: replace VN section [idx_vn+1..idx_en-1] with sorted VN
    # Easier approach: build whole new doc body

    new_doc = Document(SRC)  # start as copy
    body = new_doc.element.body

    # Remove all paragraphs (keep section properties at end)
    sectPr_list = []
    children = list(body)
    for child in children:
        if child.tag.endswith('sectPr'):
            sectPr_list.append(child)
            continue
        body.remove(child)

    # Helper: add a paragraph element (clone) to body
    def append_paragraph_xml(elem):
        # Insert before sectPr
        if sectPr_list:
            body.insert(list(body).index(sectPr_list[0]), elem)
        else:
            body.append(elem)

    # Helper: add a simple text paragraph (for headings/separators)
    def append_simple_paragraph(text, bold=False, size=None, leave_blank_before=False):
        if leave_blank_before:
            new_doc.add_paragraph('')
            # The just-added is at end of body — move before sectPr
            last = body[-1] if not sectPr_list else None
        p = new_doc.add_paragraph()
        if text:
            r = p.add_run(text)
            r.bold = bold
            if size:
                from docx.shared import Pt
                r.font.size = Pt(size)

    # Re-build body in order:
    # 1. (blank line)
    # 2. TÀI LIỆU THAM KHẢO (heading)
    # 3. (blank)
    # 4. Tiếng Việt (boundary)
    # 5. (blank)
    # 6. VN entries (sorted, with original formatting preserved)
    # 7. (blank)
    # 8. Tiếng Anh (boundary)
    # 9. (blank)
    # 10. EN entries (sorted, with original formatting preserved)

    # Easiest: clone original boundary paragraphs (idx_tlth, idx_vn, idx_en) to preserve their styling
    src_doc = Document(SRC)
    src_paras = src_doc.paragraphs

    # 1. Blank
    new_doc.add_paragraph('')
    # 2. TÀI LIỆU THAM KHẢO — clone from source
    if idx_tlth is not None:
        elem = deepcopy(src_paras[idx_tlth]._element)
        append_paragraph_xml(elem)
    # 3. Blank
    new_doc.add_paragraph('')
    # 4. Tiếng Việt — clone
    elem = deepcopy(src_paras[idx_vn]._element)
    append_paragraph_xml(elem)
    # 5. Blank
    new_doc.add_paragraph('')
    # 6. VN entries
    for sk, elem in vn_entries:
        append_paragraph_xml(elem)
    # 7. Blank
    new_doc.add_paragraph('')
    # 8. Tiếng Anh — clone
    elem = deepcopy(src_paras[idx_en]._element)
    append_paragraph_xml(elem)
    # 9. Blank
    new_doc.add_paragraph('')
    # 10. EN entries
    for sk, elem in en_entries:
        append_paragraph_xml(elem)

    new_doc.save(OUT)
    print(f'Saved: {OUT}')
    print(f'Size: {OUT.stat().st_size//1024} KB')
    print(f'Total paragraphs in output: {len(new_doc.paragraphs)}')

    # Show first 5 VN + first 5 EN sorted
    print()
    print('=== First 5 VN entries (sorted) ===')
    for sk, _ in vn_entries[:5]:
        print(f'  [{sk[:60]}]')
    print()
    print('=== First 5 EN entries (sorted) ===')
    for sk, _ in en_entries[:5]:
        print(f'  [{sk[:60]}]')

if __name__ == '__main__':
    main()
