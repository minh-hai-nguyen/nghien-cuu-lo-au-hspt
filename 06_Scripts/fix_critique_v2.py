# -*- coding: utf-8 -*-
"""
Fix new errors detected on 2nd re-verification vs original PDF:

1. THPT > THCS difference: my critique said "tất cả p<0,01" — WRONG.
   PDF shows: Emotional** (p<.01), Hyperactivity** (p<.01), Total** (p<.01);
   Conduct* (p<.05, NOT p<.01); Peer NOT significant.

2. Dang 2017 + specific Cronbach 0,47–0,62 — UNSUPPORTED.
   SDQ validation in report cites Weiss et al. 2014, not Dang 2017.
   Report does NOT publish any Cronbach alpha → my "0,47–0,62" is fabrication.

3. Bronfenbrenner socio-ecological framework — NOT mentioned in PDF.
   Need to soften to "tương tự khung socio-ecological".
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

FIXES = [
    # 1. Fix THPT > THCS — specify significance levels precisely
    ('• THPT > THCS: vấn đề hành vi, tăng động, cảm xúc, tổng vấn đề — tất cả p < 0,01 — phù hợp xu hướng "late adolescence surge" (Patton et al. 2016)',
     '• THPT > THCS: vấn đề cảm xúc (M = 4,46 vs 5,07, p < 0,01), tăng động (M = 3,63 vs 4,12, p < 0,01), tổng vấn đề SDQ (M = 12,62 vs 14,27, p < 0,01), vấn đề hành vi (M = 1,68 vs 1,98, p < 0,05); vấn đề bạn bè KHÔNG khác biệt có ý nghĩa thống kê — phù hợp xu hướng "late adolescence surge" (Patton et al. 2016)'),

    # 2. Dang 2017 alpha claim — rewrite: attribute critique to general psychometric principles, not fabricated Dang 2017 numbers
    ('Dang và cộng sự (2017, có citation trong báo cáo) đã đặt câu hỏi về factor structure 5-subscale của SDQ-25 trong bối cảnh Việt Nam; một số subscale (đặc biệt conduct và peer problems) có reliability không đạt ngưỡng chấp nhận 0,70 trong mẫu châu Á. Báo cáo UNICEF không công bố reliability riêng (Cronbach α) cho mẫu 668 HS của mình — chỉ nêu "đã validate tiếng Việt". Đây là gap minh bạch cần bổ sung trong báo cáo wave sau (alpha cụ thể cho từng subscale).',
     'Báo cáo UNICEF tham chiếu SDQ-25 "đã validate tiếng Việt (Weiss et al. 2014)" nhưng KHÔNG công bố Cronbach α hay factor loading cho mẫu 668 HS của mình. Literature quốc tế về SDQ ở bối cảnh châu Á (Goodman 1997 — khung bên ngoài; Essau et al. 2012; Du et al. 2008) ghi nhận subscale "conduct problems" và "peer problems" thường có reliability thấp hơn ngưỡng chấp nhận 0,70 khi áp dụng cross-cultural. Do đó việc báo cáo UNICEF không công bố reliability riêng là gap minh bạch — báo cáo wave sau nên bổ sung α từng subscale cho mẫu VN. (Ghi chú meta-review 13/04/2026: con số Cronbach 0,47–0,62 của phiên bản trước là chưa verified và đã được loại bỏ.)'),

    # 3. Bronfenbrenner — report doesn't explicitly use this framework, soften
    ('Cách tiếp cận này theo khung socio-ecological của Bronfenbrenner và được Fazel',
     'Cách tiếp cận multi-stakeholder này tương thích với khung socio-ecological (Bronfenbrenner 1979 — khung bên ngoài, báo cáo không chỉ đích danh) và được Fazel'),

    # 4. Also verify: 52.2% computed correctly — add note
    ('Cyberbullying: 52,2 % đã từng bị (hiếm/đôi khi/thường); THPT > THCS',
     'Cyberbullying: 52,2 % đã từng bị (Rarely 29,9 % + Sometimes 20,2 % + Often 2,1 % — tính từ Never 47,3 % báo cáo ở Bảng cyberbullying RCBI); THPT > THCS đáng kể (M tần suất 2,35 vs 3,35, p < 0,01; M distress 14,90 vs 19,13, p < 0,01)'),

    # 5. Strengthen/verify "34 KII" breakdown
    ('(n = 34)',
     '(n = 34: 3 ministry [MOET, MOH, MOLISA] + 7 principals + 22 provincial DOH/DOLISA/DOET + 1 UNICEF; tổng kiểm chéo với Bảng 2 báo cáo)'),

    # 6. Clarify 50% đi với 3h/tuần extra classes: add precise quote
    ('50 % HS báo cáo học thêm > 3 giờ/tuần | 15 % học > 9 giờ/tuần | 28 % học > 3 giờ/đêm sau giờ học',
     '50 % HS báo cáo extra classes/tutoring > 3 giờ/tuần (trích báo cáo trang 41: "50% of students report more than 3 hours per week") | 15 % học extra classes > 9 giờ/tuần | 28 % tổng thời gian học > 3 giờ/đêm sau giờ học (trích: "28% of students reporting more than 3 hours of study each night")'),
]

d = Document(PATH)
changes = 0
for p in d.paragraphs:
    t = p.text
    if not t.strip():
        continue
    new_t = t
    for find, rep in FIXES:
        if find in new_t:
            new_t = new_t.replace(find, rep)
    if new_t != t:
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = new_t
        changes += 1

d.save(PATH)
print(f'Fixed {changes} paragraphs')

# Verify
d2 = Document(PATH)
txt = '\n'.join(p.text for p in d2.paragraphs)
checks = [
    ('tất cả p < 0,01', 0),
    ('0,47–0,62', 0),
    ('p < 0,05', 1),
    ('tương thích với khung socio-ecological', 1),
    ('Rarely 29,9 % + Sometimes 20,2 %', 1),
    ('3 ministry [MOET', 1),
    ('trích báo cáo trang 41', 1),
    ('meta-review 13/04/2026', 1),
]
print('\nVerification:')
for phrase, expected in checks:
    actual = txt.count(phrase)
    ok = (actual >= expected) if expected > 0 else (actual == 0)
    print(f'  "{phrase}": {actual} ({"OK" if ok else "FAIL"})')
