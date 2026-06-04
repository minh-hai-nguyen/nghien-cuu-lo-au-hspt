"""
STEP 2: Canonical 3 bài VN + build Jenkins summary.
- VN004 NguyenThiVan 2020 STAI TPHCM
- VN005 TranNguyenNgoc 2018 Luận án GAD thư giãn
- VN006 TranThiMyLuong 2020 DASS-42 THPT Chuyên

Key facts lấy từ memory `project_cong_su_2020_papers.md` (đã verify từ cộng sự).
PDF renamed + summary docx tạo trong Tom-tat-tung-bai/.
"""
import sys, io, os, json, shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
VN_PDF_DIR = BASE / '02_Papers-goc' / 'Viet-Nam'
TT_DIR = BASE / 'Tom-tat-tung-bai'

# Định nghĩa 3 bài
PAPERS = [
    {
        'id': 'VN004',
        'descriptor': 'NguyenThiVan_2020_STAI_TPHCM',
        'pdf_src': 'NguyenThiVan_2020_STAI_TPHCM.pdf',
        'pdf_dst': 'VN004_NguyenThiVan_2020_STAI_TPHCM.pdf',
        'title_vn': 'Một số yếu tố ảnh hưởng đến rối loạn lo âu ở học sinh trung học phổ thông tại Thành phố Hồ Chí Minh',
        'title_en': 'Some causes of anxiety disorders in high school students in Ho Chi Minh City',
        'authors': 'Nguyễn Thị Vân',
        'affiliation': 'Đại học Khoa học Xã hội và Nhân văn — ĐHQG TPHCM',
        'journal': 'Tạp chí Khoa học Quản lý Giáo dục, số 01(25), tháng 3/2020. ISSN 2354-0788',
        'year': 2020,
        'sample': 'Sàng lọc 558 HS THPT (STAI) → 90 HS phân tích sâu (mẫu thuận tiện, 4 trường nội + ngoại thành TPHCM)',
        'location': 'TPHCM (4 trường THPT)',
        'tool': 'STAI Spielberger (bản Việt hoá Nguyễn Công Khanh 2000) + SAS',
        'design': 'Cắt ngang, 2 pha (sàng lọc + khảo sát sâu)',
        'outcomes': [
            'Tỷ lệ RLLA: 15–18,5 % (STAI ≥ ngưỡng)',
            '4 nhóm yếu tố: học tập / gia đình / quan hệ xã hội / bản thân HS',
            'Nhóm "bản thân HS" ảnh hưởng mạnh nhất: r = 0,42 (tác giả công bố 85,4 % — em nghi là R² hồi quy đa biến nhiều biến bản thân, không phải r² bivariate vì r²=0,176)',
            'Top 3 yếu tố bản thân: ít giao tiếp chia sẻ (X̄=0,91), thất vọng (0,90), lo sợ khó khăn (0,88)',
            'Nhóm học tập: r = 0,37; áp lực thi ĐH 56,7 %; định hướng nghề 51,5 %; kỳ vọng cha mẹ 48,9 %',
        ],
        'strengths': [
            'Một trong số ít NC VN chuyên về YẾU TỐ ẢNH HƯỞNG (không chỉ prevalence)',
            'Phân tích 4 nhóm yếu tố hệ thống',
            'Kết hợp số liệu + phỏng vấn định tính',
        ],
        'limitations': [
            'Mẫu thuận tiện (không ngẫu nhiên) → không đại diện TPHCM',
            'STAI bản Việt 2000 — chưa cập nhật psychometric trong 20 năm',
            'Thiếu phân tích đa biến rõ ràng (R² 85,4 % cần verify)',
            'Thiếu nhóm chứng',
            'Chỉ 1 thành phố, không phân biệt quận đô thị vs ngoại thành sâu',
        ],
        'critique_short': 'NC quan trọng về YẾU TỐ nguy cơ RLLA ở HS THPT VN (hiếm có). Tuy nhiên mẫu thuận tiện + công cụ cũ + số liệu R² gây nghi ngờ. Con số 85,4 % giải thích của nhóm "bản thân HS" không khớp với r=0,42 (r²=0,176) — có thể là R² đa biến hoặc tỷ lệ % nhóm có biểu hiện. Cần verify bản gốc.',
        'future': 'Sao chép thiết kế với mẫu ngẫu nhiên đại diện + STAI cập nhật hoặc DASS-21 + phân tích đa biến rõ ràng + nhóm chứng.',
        'rating': 3,
        'notes_for_rag': 'Trong 4 nhóm yếu tố: "bản thân HS" (personality/self) MẠNH NHẤT > học tập > gia đình > quan hệ xã hội.',
    },
    {
        'id': 'VN005',
        'descriptor': 'TranNguyenNgoc_2018_LuanAn_ThuGian_GAD',
        'pdf_src': 'TranNguyenNgoc_2018_LuanAn_TSYH_BachMai.pdf',
        'pdf_dst': 'VN005_TranNguyenNgoc_2018_LuanAn_ThuGian_GAD.pdf',
        'title_vn': 'Đánh giá hiệu quả điều trị rối loạn lo âu lan tỏa bằng liệu pháp thư giãn luyện tập',
        'title_en': 'Evaluation of treatment efficacy for generalized anxiety disorder using relaxation training therapy',
        'authors': 'Trần Nguyễn Ngọc',
        'affiliation': 'Đại học Y Hà Nội / Bạch Mai',
        'journal': 'Luận án Tiến sĩ Y học, ĐH Y Hà Nội, 2018 (177 trang)',
        'year': 2018,
        'sample': 'Bệnh nhân GAD (chẩn đoán DSM/ICD) điều trị tại bệnh viện — số liệu cụ thể cần verify từ chương "Đối tượng và phương pháp"',
        'location': 'Bệnh viện Bạch Mai, Hà Nội',
        'tool': 'DSM/ICD chẩn đoán + thang đo lo âu (HAM-A, Beck, Zung — cần verify) + liệu pháp thư giãn luyện tập',
        'design': 'Can thiệp trước-sau (có thể có nhóm chứng) — thiết kế cụ thể verify từ PDF',
        'outcomes': [
            'Công trình nghiên cứu CAN THIỆP RLLA tại Việt Nam — RẤT HIẾM trong thư viện',
            'Trọng tâm: liệu pháp thư giãn luyện tập (relaxation training) cho GAD người lớn',
            'Đánh giá hiệu quả giảm triệu chứng trước-sau can thiệp (số cụ thể cần verify)',
        ],
        'strengths': [
            'Công trình VN về can thiệp RLLA — hầu như không có công trình khác cùng hướng trong thư viện',
            'Luận án TS → methodology được bảo vệ',
            'Bệnh viện Bạch Mai — trung tâm tâm thần học đầu ngành VN',
            '177 trang chi tiết (hiếm có ở VN)',
        ],
        'limitations': [
            'Đối tượng NGƯỜI LỚN (không phải HS THCS/THPT — trọng tâm đề tài) — giá trị liên hệ gián tiếp',
            'Liệu pháp thư giãn đơn thuần, không phải CBT đầy đủ',
            'Chưa publish quốc tế — visibility thấp',
            'Năm 2018 — dữ liệu có thể outdated với bối cảnh hậu COVID',
        ],
        'critique_short': 'Công trình can thiệp RLLA HIẾM ở VN — có giá trị làm nền cho RCT can thiệp HS/VTN tương lai. Nhưng đối tượng là NGƯỜI LỚN, liệu pháp đơn chỉ thư giãn — không thay thế CBT hoặc kết hợp. Dùng làm tài liệu tham chiếu methodology, không dùng cho effect size HS.',
        'future': 'Mở rộng thiết kế: (1) đối tượng VTN 12–17 tuổi; (2) so sánh thư giãn vs CBT nhóm vs kết hợp; (3) RCT với randomization + nhóm chứng chặt chẽ.',
        'rating': 3,
        'notes_for_rag': 'Can thiệp RLLA tại VN — thư giãn luyện tập, bệnh viện Bạch Mai, người lớn, 2018, luận án TS.',
    },
    {
        'id': 'VN006',
        'descriptor': 'TranThiMyLuong_2020_DASS42_THPTChuyen',
        'pdf_src': 'TranThiMyLuong_2020_DASS42_THPTChuyen.pdf',
        'pdf_dst': 'VN006_TranThiMyLuong_2020_DASS42_THPTChuyen.pdf',
        'title_vn': 'Rối loạn lo âu ở học sinh trung học phổ thông chuyên',
        'title_en': 'Anxiety disorders in students of specialized (selective) high schools',
        'authors': 'Trần Thị Mỵ Lương (và có thể Đặng Đức Anh)',
        'affiliation': 'Học viện Phụ nữ Việt Nam',
        'journal': 'Tạp chí Khoa học ĐH Thủ Đô Hà Nội, số 40/2020, tr 122',
        'year': 2020,
        'sample': '540 HS THPT chuyên → 77 HS có RLLA (tỷ lệ 14,2 %)',
        'location': 'Trường THPT chuyên (1 trường, cần verify tên trường cụ thể từ PDF)',
        'tool': 'DASS-42 (40 mục về lo âu — cutoff không nêu rõ trong bản tóm tắt)',
        'design': 'Cắt ngang, sàng lọc',
        'outcomes': [
            'Tỷ lệ lo âu tổng: 14,2 % (nhẹ 3,5 % / vừa 7,2 % / nặng 2,4 % / rất nặng 1,1 %)',
            'PHÁT HIỆN BẤT NGỜ: Khối 11 chiếm 48,1 % (CAO NHẤT), khối 10 31,1 %, khối 12 chỉ 20,8 %',
            '→ TRÁI giả thuyết rằng khối 12 (thi ĐH) cao nhất. Lý giải: khối 11 áp lực đa chiều chọn ngành/khối; khối 12 đã ổn định định hướng',
            'Biểu hiện: choáng váng 85,7 %, khó tập trung 76,6 %, cảm thấy "vô dụng" 67,5 %',
            'Trích phỏng vấn: "Em là HS Chuyên mà mấy cái đó cũng không làm được, em thấy tệ quá!" → "perfectionism trap" trường chuyên',
            'Giới: nữ 66,5 % (lệch giới)',
        ],
        'strengths': [
            'Một trong ít NC VN chuyên về HS TRƯỜNG CHUYÊN — đặc thù áp lực cao',
            'Phân tầng theo khối lớp → phát hiện mẫu hình TRÁI với giả thuyết',
            'Kết hợp định lượng + phỏng vấn định tính',
            'DASS-42 (full version) thay vì chỉ DASS-21 → thêm chi tiết',
        ],
        'limitations': [
            'Chỉ 1 trường chuyên đặc thù → không ngoại suy ra toàn bộ HS chuyên VN',
            'Không nêu cutoff DASS-42 cụ thể → khó đối chiếu với NC khác',
            'Lệch giới (nữ 66,5 %) → có thể phóng đại tỷ lệ',
            'Thiếu phân tích đa biến',
            'Thiếu nhóm chứng (so với HS không-chuyên)',
        ],
        'critique_short': 'Phát hiện KHỐI 11 > KHỐI 12 lo âu trong trường chuyên là đóng góp đặc sắc, ngược giả định "thi ĐH gây lo âu nhất". Gợi ý cho can thiệp: tập trung vào KHỐI 11 trường chuyên, không phải khối 12. Tuy nhiên chỉ 1 trường nên cần nhân rộng để confirm.',
        'future': 'So sánh HS chuyên vs HS không-chuyên cùng khối với cùng công cụ. Điều tra yếu tố "perfectionism" qua thang đo MPS. Theo dõi dọc khối 11 → khối 12 xem lo âu có giảm tự nhiên không.',
        'rating': 4,
        'notes_for_rag': 'HS THPT chuyên khối 11 lo âu cao nhất 48,1 %, TRÁI với giả định khối 12 cao nhất. Perfectionism trap. DASS-42 n=540.',
    },
]

# =============================================================
# Build docx Jenkins summary
# =============================================================
def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def build_summary_docx(paper, out_path):
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # Title
    t = d.add_heading(f"{paper['id']} — Tóm tắt bài nghiên cứu", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)

    # Metadata table
    meta_tbl = d.add_table(rows=7, cols=2)
    meta_tbl.style = 'Table Grid'
    meta_rows = [
        ('Tên công trình (VN)', paper['title_vn']),
        ('Tên công trình (EN)', paper['title_en']),
        ('Tác giả', paper['authors']),
        ('Đơn vị', paper['affiliation']),
        ('Nơi công bố', paper['journal']),
        ('Năm', str(paper['year'])),
        ('Địa bàn + mẫu + công cụ', f"{paper['location']} | {paper['sample']} | Công cụ: {paper['tool']}"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        cell0 = meta_tbl.rows[i].cells[0]
        shade_cell(cell0, 'D9E1F2')
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        meta_tbl.rows[i].cells[1].text = v
    d.add_paragraph()

    # Phương pháp
    h = d.add_heading('Phương pháp nghiên cứu', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)
    d.add_paragraph(paper['design'])
    d.add_paragraph(f"Mẫu: {paper['sample']}")
    d.add_paragraph(f"Công cụ đo: {paper['tool']}")

    # Kết quả
    h = d.add_heading('Kết quả nghiên cứu', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)
    for it in paper['outcomes']:
        d.add_paragraph(it, style='List Bullet')

    # Nhận xét — điểm mạnh
    h = d.add_heading('Điểm mạnh', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(54, 95, 44)
    for it in paper['strengths']:
        d.add_paragraph('• ' + it)

    # Hạn chế
    h = d.add_heading('Hạn chế', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(192, 80, 77)
    for it in paper['limitations']:
        d.add_paragraph('• ' + it)

    # Phản biện ngắn
    h = d.add_heading('Phản biện ngắn', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(192, 0, 0)
    p = d.add_paragraph()
    r = p.add_run(paper['critique_short'])
    r.font.color.rgb = RGBColor(192, 0, 0)

    # Hướng nghiên cứu tiếp
    h = d.add_heading('Hướng nghiên cứu tiếp theo', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)
    d.add_paragraph(paper['future'])

    # Đánh giá
    stars = '⭐' * paper['rating']
    h = d.add_heading(f"Đánh giá chất lượng: {stars} ({paper['rating']}/5)", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(191, 97, 14)

    # Ghi chú cho RAG
    if paper.get('notes_for_rag'):
        d.add_paragraph()
        p = d.add_paragraph()
        r = p.add_run('Key fact cho RAG: ')
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(90, 90, 90)
        p.add_run(paper['notes_for_rag']).font.size = Pt(10)

    d.save(out_path)

# =============================================================
# Execute
# =============================================================
log = []

with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

for paper in PAPERS:
    # 1. Rename PDF
    src = VN_PDF_DIR / paper['pdf_src']
    dst = VN_PDF_DIR / paper['pdf_dst']
    if not src.exists():
        log.append(f"SKIP {paper['id']}: PDF source missing {paper['pdf_src']}")
        continue
    if not dst.exists():
        src.rename(dst)
        log.append(f"RENAMED {paper['id']}: {src.name} → {dst.name}")
    else:
        log.append(f"SKIP rename {paper['id']}: dst already exists")

    # 2. Build summary docx
    tt_name = f"{paper['id']}_{paper['descriptor']}.docx"
    tt_path = TT_DIR / tt_name
    build_summary_docx(paper, tt_path)
    log.append(f"SUMMARY {paper['id']}: {tt_name} ({tt_path.stat().st_size//1024} KB)")

    # 3. Update canonical_index
    idx[paper['id']] = {
        'id': paper['id'],
        'descriptor': paper['descriptor'],
        'summary': tt_name,
        'translation': None,
        'pdf': paper['pdf_dst'],
        'pdf_folder': 'Viet-Nam',
    }
    log.append(f"INDEX {paper['id']}: added to canonical_index.json")

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

log.append(f"\\nTotal canonical entries now: {len(idx)}")
log.append(f"VN entries: {sum(1 for k in idx if k.startswith('VN'))}")

for l in log:
    print(l)
