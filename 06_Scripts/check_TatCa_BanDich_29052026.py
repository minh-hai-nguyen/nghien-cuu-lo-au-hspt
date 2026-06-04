# -*- coding: utf-8 -*-
"""Bo test xuyen-file: kiem tra dich va consistency cua TAT CA cac ban dich
+ LA chinh, theo tieu chuan 15-category framework da phat trien.

KHONG SUA FILE, chi RECORD loi de bao cao.
29/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

FILES = {
    'LA': 'Luận án TS/01_LuanAn_v4_ChuanTrinhBay_28052026.docx',
    'TT_VN': 'Luận án TS/TomTatLA_v2_VERIFIED_29052026.docx',
    'TT_EN': 'Luận án TS/TomTatLA_EN_v1_29052026.docx',
    'TY': 'Luận án TS/TrichYeuLA_CongThiHang_v2_29052026.docx',
    'DM': 'Luận án TS/DanhMucCongTrinh_EN_v1_29052026.docx',
    'BaiA': 'bai-bao-Q1/BaiA_JAD_OUTLINE_v1_30052026.docx',
    'BaiD': 'bai-bao-Q1/BaiD_StressHealth_OUTLINE_v1_30052026.docx',
}

texts = {}
for key, path in FILES.items():
    full = os.path.join(ROOT, path)
    if not os.path.exists(full):
        print(f'WARNING: {key} not found at {full}')
        texts[key] = ''
        continue
    d = Document(full)
    t = '\n'.join(p.text for p in d.paragraphs)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                t += '\n' + c.text
    texts[key] = t

all_errors = []


def report(file_key, category, name, ok, detail=''):
    """Record check result. Print + add to error log if fail."""
    status = '✓' if ok else '✗'
    print(f'  [{file_key}] {status} {category}/{name}')
    if detail:
        print(f'        {detail}')
    if not ok:
        all_errors.append(f'[{file_key}] {category}/{name}' + (f' | {detail}' if detail else ''))


# ============================================================
# CATEGORY A: GOLDEN FACTS (from LA chinh)
# ============================================================
# These are the source-of-truth facts from LA. ALL files (where relevant) should match.
GOLDEN = {
    'so_hs_vn': '1.352',
    'so_hs_en': '1,352',
    'tieu_de_vn_part': 'rối loạn lo âu ở học sinh',
    'ma_so': '9.31.04.01',
    'F_chia_ly_vn': '0,246',
    'F_chia_ly_en': '0.246',
    'p_chia_ly_vn': '0,620',
    'p_chia_ly_en': '0.620',
    'tt_pct_range_vn': '85–89',
    'dtb_alht': '51,13',
    'dtb_alht_en': '51.13',
    'dtb_ndt': '28,38',
    'dtb_bnhd': '13,52',
    'khung_ct_8nd': '8 nội dung',
    'ncs_vn': 'Công Thị Hằng',
    'ncs_en': 'Cong Thi Hang',
    'cbhd_vn': 'TS. Đào Minh Đức',
    'cbhd_en_western': 'Hang Thi Cong',  # for Q1 papers only
    'cbhd_en_vn_order': 'Dr. Dao Minh Duc',  # for LA + TY (Vietnamese order)
}

# Scale acronyms (8 - verified in LA)
SCALES = {
    'RCADS': 'Chorpita',
    'ESSA': 'Sun',
    'SAS-SV': 'Kwon',
    'OBVQ': 'Olweus',
    'PSSM': 'Goodenow',
    'MSPSS': 'Zimet',
    'RSES': 'Rosenberg',
    'Brief COPE': 'Carver',
}

# FABRICATED names that should NOT appear in ANY file
FABRICATED = ['MPAS', 'MPVS']

# OLD/STALE values that should NOT appear
STALE = {
    '~86%': 'cường độ TTr cũ sai',
    '9310401': 'mã số format cũ (template Khôi)',
    'lo âu tổng quát': 'thuật ngữ sai (DSM-5 dùng "lan tỏa")',
    'junior secondary school': 'thuật ngữ sai (MOET dùng "lower secondary")',
}


# ============================================================
# CHECK 1: LA CHINH internal consistency
# ============================================================
print('\n=== CHECK 1: LA CHÍNH internal consistency ===')
la = texts['LA']

# All 8 scales present
for scale in SCALES:
    cnt = la.count(scale)
    report('LA', 'scale', f'"{scale}" >= 5 lần', cnt >= 5, f'count={cnt}')

# Golden facts in LA
for k, v in [('1.352', '1.352'), ('9.31.04.01', '9.31.04.01'),
             ('51,13', '51,13'), ('28,38', '28,38'), ('13,52', '13,52'),
             ('85–89', 'TTr range'), ('8 nội dung', 'Khung CT'),
             ('F = 0,246', 'F sep'), ('p = 0,620', 'p sep')]:
    report('LA', 'fact', k, k in la, '')

# No fabricated names
for fab in FABRICATED:
    report('LA', 'anti-fab', f'KHÔNG "{fab}"', fab not in la)


# ============================================================
# CHECK 2: TT_VN (Tóm tắt LA Vietnamese)
# ============================================================
print('\n=== CHECK 2: TÓM TẮT LA VN (v2) ===')
tv = texts['TT_VN']

# Match LA on key numbers
for k in ['1.352', '9.31.04.01', '51,13', '28,38', '13,52',
         '85–89', '8 nội dung', 'F = 0,246', 'p = 0,620']:
    report('TT_VN', 'fact_match_LA', k, k in tv)

# All scales present
for scale in SCALES:
    report('TT_VN', 'scale', scale, scale in tv)

# Anti-fab
for fab in FABRICATED:
    report('TT_VN', 'anti-fab', f'no "{fab}"', fab not in tv)
for stale, desc in STALE.items():
    report('TT_VN', 'anti-stale', f'no "{stale}" ({desc})', stale not in tv)

# Author name VN order
report('TT_VN', 'name', 'Công Thị Hằng', 'Công Thị Hằng' in tv)
report('TT_VN', 'name', 'TS. Đào Minh Đức', 'TS. Đào Minh Đức' in tv)


# ============================================================
# CHECK 3: TT_EN (Tóm tắt LA English)
# ============================================================
print('\n=== CHECK 3: TÓM TẮT LA EN (v1) ===')
te = texts['TT_EN']

# Numbers in EN format
for k in ['1,352', '9.31.04.01']:
    report('TT_EN', 'fact_match', k, k in te)

# Scales
for scale in SCALES:
    report('TT_EN', 'scale', scale, scale in te)

# Anti-fab + anti-stale
for fab in FABRICATED:
    report('TT_EN', 'anti-fab', f'no "{fab}"', fab not in te)

# Author EN — for tóm tắt LA (official), uses Vietnamese order (Cong Thi Hang)
# but earlier we discussed potentially using Western order — let me check what's in file
has_vn_order = 'Cong Thi Hang' in te
has_western = 'Hang Thi Cong' in te
report('TT_EN', 'name', 'Author EN format (VN or Western order)',
       has_vn_order or has_western,
       f'VN_order={has_vn_order} | Western={has_western}')

# Statistical format - EN uses period
report('TT_EN', 'stat', 'F = 0.246 (period)', 'F = 0.246' in te or 'F=0.246' in te)
report('TT_EN', 'stat', 'p = 0.620 (period)', 'p = 0.620' in te or 'p=0.620' in te)


# ============================================================
# CHECK 4: TY (Trích yếu LA v2)
# ============================================================
print('\n=== CHECK 4: TRÍCH YẾU LA (v2) ===')
ty = texts['TY']

# Already 206 checks PASS — sanity-check key items
for k in ['1.352', '9.31.04.01', '85–89', '8 nội dung',
         'F = 0,246', 'F = 0.246', 'p = 0,620', 'p = 0.620']:
    report('TY', 'fact', k, k in ty)

for scale in SCALES:
    report('TY', 'scale', scale, scale in ty)

for fab in FABRICATED:
    report('TY', 'anti-fab', f'no "{fab}"', fab not in ty)
for stale, _ in STALE.items():
    report('TY', 'anti-stale', f'no "{stale}"', stale not in ty)

report('TY', 'name', 'Công Thị Hằng (VN)', 'Công Thị Hằng' in ty)
report('TY', 'name', 'Cong Thi Hang (EN VN order)', 'Cong Thi Hang' in ty)
report('TY', 'name', 'Dr. Dao Minh Duc', 'Dr. Dao Minh Duc' in ty)


# ============================================================
# CHECK 5: DM (Danh mục công trình EN với 2 ref Q2.5)
# ============================================================
print('\n=== CHECK 5: DANH MỤC CÔNG TRÌNH EN ===')
dm = texts['DM']

# 2 references Q2.5 - format Western order
report('DM', 'name', '"Hang Thi Cong" Western order',
       'Hang Thi Cong' in dm)

# Journal names corrected
report('DM', 'journal', 'Vietnam Journal of Educational Sciences',
       'Vietnam Journal of Educational Sciences' in dm)
report('DM', 'journal', 'Vietnam Journal of Psychology',
       'Vietnam Journal of Psychology' in dm)

# No "Review" (sai)
review_count = dm.count('Review Khoa') + dm.count('Review Tâm lý')
report('DM', 'journal', 'KHÔNG còn "Review Khoa.../Tâm lý..." (sai cũ)',
       review_count == 0, f'count={review_count}')

# No "junior secondary school"
report('DM', 'terminology', 'no "junior secondary"',
       'junior secondary' not in dm.lower())

# Has "lower secondary school"
report('DM', 'terminology', '"lower secondary school" present',
       'lower secondary school' in dm.lower())

# "May" capitalized
report('DM', 'format', '"May" capitalized (not "may")',
       'May 2026' in dm or 'May)' in dm)


# ============================================================
# CHECK 6: BaiA (Outline Q1 JAD)
# ============================================================
print('\n=== CHECK 6: OUTLINE BÀI A Q1 JAD ===')
ba = texts['BaiA']

# Numbers consistency with LA
report('BaiA', 'fact', '1,352 students', '1,352' in ba or '1.352' in ba)

# Scales mentioned (should match LA)
for scale in SCALES:
    has = scale in ba
    report('BaiA', 'scale', scale, has)

# Anti-fabrication
for fab in FABRICATED:
    report('BaiA', 'anti-fab', f'no "{fab}"', fab not in ba)

# Author name format - Q1 uses Western order
report('BaiA', 'name', 'Western order "Hang Thi Cong"',
       'Hang Thi Cong' in ba)

# Should NOT use Vietnamese order "Cong Thi Hang" in byline (only in cite "Cong et al")
# But "Cong" alone is OK for citation
report('BaiA', 'stat', 'p < 0.001 (English decimal)',
       'p < 0.001' in ba or 'p<0.001' in ba)


# ============================================================
# CHECK 7: BaiD (Outline Q1 Stress and Health)
# ============================================================
print('\n=== CHECK 7: OUTLINE BÀI D Q1 STRESS AND HEALTH ===')
bd = texts['BaiD']

report('BaiD', 'fact', '1,352 students', '1,352' in bd or '1.352' in bd)

# Brief COPE specifically central to Bài D
report('BaiD', 'scale_focus', 'Brief COPE central', 'Brief COPE' in bd)
report('BaiD', 'scale_focus', 'Carver (1997)',
       'Carver' in bd and '1997' in bd)

# Anti-fab
for fab in FABRICATED:
    report('BaiD', 'anti-fab', f'no "{fab}"', fab not in bd)

# Author name Western
report('BaiD', 'name', 'Western order "Hang Thi Cong"',
       'Hang Thi Cong' in bd)

# Confucian/Vietnamese cultural framing (Bài D story angle)
report('BaiD', 'frame', 'Cultural framing (nhẫn / Confucian)',
       'nhẫn' in bd.lower() or 'Confucian' in bd.lower())


# ============================================================
# CHECK 8: CROSS-FILE CONSISTENCY
# ============================================================
print('\n=== CHECK 8: CROSS-FILE CONSISTENCY ===')

# 8.1 Mã số same in TT_VN, TY (both LA-context files should match LA)
ma_so_files = {k: ('9.31.04.01' in v) for k, v in texts.items()}
for k in ['LA', 'TT_VN', 'TY']:
    report('cross', 'ma_so', f'"9.31.04.01" trong {k}', ma_so_files[k])

# 8.2 Số HS consistent
ss_vn = '1.352' in texts['LA'] and '1.352' in texts['TT_VN'] and '1.352' in texts['TY']
ss_en = '1,352' in texts['TT_EN'] and '1,352' in texts['BaiA'] and '1,352' in texts['BaiD']
report('cross', 'sample_size', '1.352 (VN) consistent', ss_vn)
report('cross', 'sample_size', '1,352 (EN) consistent', ss_en)

# 8.3 Scale names consistent across all docs that mention them
for scale in SCALES:
    docs_with = [k for k in ['LA', 'TT_VN', 'TT_EN', 'TY']
                 if scale in texts[k]]
    report('cross', 'scale_consistency',
           f'"{scale}" trong LA + TT_VN + TT_EN + TY',
           len(docs_with) == 4, f'present in: {docs_with}')

# 8.4 No fabricated names anywhere
for fab in FABRICATED:
    bad_files = [k for k, t in texts.items() if fab in t]
    report('cross', 'anti-fab',
           f'KHÔNG có "{fab}" ở bất kỳ file nào',
           len(bad_files) == 0, f'found in: {bad_files}')

# 8.5 No stale ~86% anywhere
bad_files = [k for k, t in texts.items() if '~86%' in t or 'khoảng 86%' in t]
report('cross', 'anti-stale', 'KHÔNG có "~86%" stale', len(bad_files) == 0,
       f'found in: {bad_files}')

# 8.6 No "9310401" stale (only "9.31.04.01" expected)
bad_files = [k for k, t in texts.items() if '9310401' in t]
report('cross', 'anti-stale', 'KHÔNG có "9310401" stale', len(bad_files) == 0,
       f'found in: {bad_files}')

# 8.7 Tên đề tài consistent
title_vn_present = {
    k: 'rối loạn lo âu ở học sinh' in texts[k].lower()
    for k in ['LA', 'TT_VN', 'TY']
}
all_consistent = all(title_vn_present.values())
report('cross', 'title', 'Title VN "Rối loạn lo âu ở HS THCS" nhất quán',
       all_consistent, f'detail={title_vn_present}')

# 8.8 Citation style for Q1 papers
# Western order in BaiA + BaiD + DM
western_present = {
    k: 'Hang Thi Cong' in texts[k]
    for k in ['BaiA', 'BaiD', 'DM']
}
report('cross', 'name_format_Q1',
       'Q1 files đều dùng Western order "Hang Thi Cong"',
       all(western_present.values()), f'detail={western_present}')

# 8.9 Vietnamese order for LA-context files (TT_EN, TY)
vn_order_present = {
    k: 'Cong Thi Hang' in texts[k] and 'Hang Thi Cong' not in texts[k]
    for k in ['TT_EN', 'TY']
}
report('cross', 'name_format_LA',
       'LA files dùng VN order "Cong Thi Hang" (KHÔNG có Western)',
       all(vn_order_present.values()), f'detail={vn_order_present}')


# ============================================================
# CHECK 9: SCALE-FACT ALIGNMENT
# ============================================================
print('\n=== CHECK 9: SCALE NAME ↔ AUTHOR YEAR ALIGNMENT ===')
# For each scale, where it appears, the author should appear near it
scale_year = {
    'RCADS': ('Chorpita', '2000'),
    'ESSA': ('Sun', '2011'),
    'SAS-SV': ('Kwon', '2013'),
    'OBVQ': ('Olweus', '1996'),
    'PSSM': ('Goodenow', '1993'),
    'MSPSS': ('Zimet', '1988'),
    'RSES': ('Rosenberg', '1965'),
    'Brief COPE': ('Carver', '1997'),
}

for fkey in ['TT_VN', 'TT_EN', 'TY']:
    t = texts[fkey]
    for scale, (author, year) in scale_year.items():
        if scale in t:
            has_pair = author in t and year in t
            report(fkey, 'scale_author', f'{scale}+{author}+{year}',
                   has_pair)


# ============================================================
# SUMMARY
# ============================================================
print()
print('=' * 70)
print(f'TỔNG SỐ LỖI: {len(all_errors)}')
if all_errors:
    print('\nDANH SÁCH LỖI (sẽ ghi nhớ, KHÔNG sửa):')
    for i, err in enumerate(all_errors, 1):
        print(f'  {i}. {err}')

    # Write to file for record
    log_path = os.path.join(ROOT, '06_Scripts', 'errors_log_xuyenfile_29052026.txt')
    with open(log_path, 'w', encoding='utf-8') as f:
        f.write(f'Errors found: {len(all_errors)}\n')
        f.write('=' * 70 + '\n')
        for i, err in enumerate(all_errors, 1):
            f.write(f'{i}. {err}\n')
    print(f'\n📄 Đã ghi log lỗi: {log_path}')
else:
    print('\nSẠCH LỖI XUYÊN FILE!')
