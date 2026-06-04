# -*- coding: utf-8 -*-
"""Dịch + Tóm tắt QT36 — Korea GAD ML 2025 — Moon & Woo"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

# ===== TÓM TẮT =====
doc2 = create_doc()
add_heading(doc2, 'Tóm tắt bài QT-36', 1)
add_p(doc2, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
add_p(doc2, 'Công trình \u00ab Các yếu tố nguy cơ chính của GAD ở VTN: Nghiên cứu học máy \u00bb (Key risk factors of GAD in adolescents: machine learning study), do Yonghwan Moon và Hyekyung Woo (2025), ĐH Quốc gia Kongju, Hàn Quốc, phân tích 213.820 VTN Hàn Quốc (KYRBS 2020\u20132023; THCS 53,8% + THPT 46,2%). Frontiers in Public Health, 2025. DOI: 10.3389/fpubh.2024.1504739.')

add_p(doc2, 'Phương pháp nghiên cứu', bold=True)
add_p(doc2, 'Sử dụng Machine Learning (LASSO, SelectKBest, XGBoost) để chọn đặc trưng + Random Forest và ANN để dự đoán GAD (GAD-7 \u2265 10). Nói cách khác, đây là NC ĐẦU TIÊN dùng ML đa phương pháp để xác định yếu tố nguy cơ GAD riêng ở VTN \u2014 không dùng hồi quy truyền thống.')
add_p(doc2, 'KYRBS (Korea Youth Risk Behavior Web-based Survey) 2020\u20132023, quản lý bởi KDCA. n = 213.820 VTN. THCS 53,8% + THPT 46,2%. Phân tích: 3 thuật toán chọn đặc trưng \u00d7 2 mô hình dự đoán = 6 tổ hợp. SMOTE xử lý mất cân bằng.')

add_p(doc2, 'Kết quả nghiên cứu định lượng', bold=True)

add_heading(doc2, 'Bảng 1. Top yếu tố nguy cơ GAD (ML, n = 213.820)', 3)
add_table(doc2,
    ['Hạng', 'Yếu tố', 'LASSO', 'F-score', 'XGBoost', '\u00dd nghĩa'],
    [['1', 'Cô đơn', '0,23', '16.646', '0,29', 'MẠNH NHẤT \u2014 cả 3 thuật toán'],
     ['2', 'Stress', '0,10', '6.589', '0,07', 'Nhất quán'],
     ['3', 'Sức khỏe chủ quan kém', '0,12', '6.782', '0,04', ''],
     ['4', 'Giấc ngủ phục hồi kém', '0,10', '5.081', '0,02', 'Zhu 2025 QT05: AOR=13,71'],
     ['5', 'Bạo lực', '0,10', '5.117', '0,01', 'VN15: ACEs\u2192lo âu']],
    widths=[0.8, 3.5, 1.5, 2.0, 2.0, 3.5])

add_heading(doc2, 'Bảng 2. Hiệu suất mô hình (RF + XGBoost features)', 3)
add_table(doc2,
    ['Chỉ số', 'Giá trị', '\u00dd nghĩa'],
    [['Accuracy', '78%', ''],
     ['AUC', '82%', 'Tốt'],
     ['Sensitivity', '64%', 'Trung bình \u2014 36% bỏ sót'],
     ['Prevalence GAD', '17,9%', 'GAD-7 \u2265 10']],
    widths=[3.0, 2.0, 5.0])

add_p(doc2, 'Đối chiếu liên bài', bold=True)
add_p(doc2, 'Cô đơn hạng 1 \u2014 phù hợp Ireland 2024 (QT32, OGA bảo vệ), Islam 2025 (QT31, không bạn thân AOR=1,28). Giấc ngủ \u2014 phù hợp Zhu 2025 (QT05, <5h AOR=13,71). GAD-7 \u2014 phù hợp Hoa 2024 (VN01, 40,6% nhưng \u22655). ML trên KYRBS \u2014 phù hợp Korea 2024 (QT34, KYRBS 16 năm).')

add_p(doc2, 'Nhận xét, phát hiện qua kết quả nghiên cứu', bold=True)
add_p(doc2, '*Cô đơn vượt trội stress.* Gợi ý: can thiệp GIẢM CÔ ĐƠN (kết nối bạn bè, mentor/OGA) có thể hiệu quả hơn can thiệp stress đơn thuần. *Hành vi sức khỏe (giấc ngủ, hút thuốc, dinh dưỡng) \u2014 CAN THIỆP ĐƯỢC.* *ML phát hiện quan hệ phi tuyến* mà hồi quy truyền thống bỏ sót.')

add_p(doc2, 'Kết luận', bold=True)
add_p(doc2, 'Dữ liệu 213.820 VTN Hàn Quốc với ML, cho thấy cô đơn là yếu tố mạnh nhất liên quan GAD (F=16.646), tiếp theo stress và sức khỏe chủ quan, gợi ý rằng chiến lược giảm cô đơn (kết nối bạn bè, mentor) + cải thiện giấc ngủ + phòng hút thuốc là ưu tiên can thiệp phòng ngừa GAD ở VTN.')

add_red_heading(doc2, 'Phản biện')
add_red(doc2, 'Frontiers Q1 IF=5,2. ML đa phương pháp, n=213.820, GAD-7 riêng. Nhưng: chỉ Hàn Quốc, cắt ngang, SMOTE overfitting, sensitivity 64%, thiếu MXH/áp lực HT/ACEs.')

add_red_heading(doc2, 'Hướng NC tiếp theo')
add_red(doc2, 'ML trên dữ liệu VN (V-NAMHS, GSHS). So sánh KYRBS vs GSHS VN. NC cô đơn ở VTN VN. Kết hợp ML + NC dọc (Hue Cohort VN21).')

add_p(doc2, 'Đánh giá: \u2b50\u2b50\u2b50\u2b50 Cao. Frontiers Q1, ML mới, n=213.820, GAD-7 riêng, phát hiện cô đơn hạng 1.', bold=True)

doc2.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'Tom-tat-tung-bai', 'QT36_Korea_GAD_ML_2025.docx'))
print('Tom tat QT36 saved!')

import docx as dx
d2 = dx.Document(os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'Tom-tat-tung-bai', 'QT36_Korea_GAD_ML_2025.docx'))
t2 = '\n'.join([p.text for p in d2.paragraphs])
print(f'  Chars: {len(t2)}, Tables: {len(d2.tables)}')
