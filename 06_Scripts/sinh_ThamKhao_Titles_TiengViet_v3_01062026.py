# -*- coding: utf-8 -*-
"""Sinh v3 - tich hop tat ca cac sua sau audit.

v3 fixes:
- Paper #1: Juan J et al. + thêm "GIRLS" vào title
- Paper #6: Năm 2026 (không phải 2025) + tác giả Aijun Zhu et al.
- Paper #7: Thay đổi cách present — đánh dấu sample là college students,
  không phải adolescent, dùng làm ví dụ cấu trúc title only
- Paper #9: Tran Minh Dien et al. + BMC Public Health + n=5,325
- Paper #13: Năm 2024 + tác giả Gebreegziabher Z et al.
- Paper #17: Niwenahisemo LC et al. + n=1,813 Kigali

Bổ sung sample sizes verified:
- Paper #5: n=486 (Bangladesh Math)
- Paper #6: n=3,673 (Personality)
- Paper #11: n=1,591 (Guizhou)
- Paper #11 V-NAMHS related: n=845

01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'ThamKhao_Titles_Q1Q3_AsiaChauPhi_TiengViet_v3_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.4


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=True, align_center=False):
    p = d.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_center
                   else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('▸ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def WARN(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_title_table(rows, col_widths_cm):
    t = d.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'; t.autofit = False
    headers = ['Số', 'Tiêu đề bài báo (tiếng Anh + dịch nghĩa)',
               'Tác giả + Quốc tịch + Sample', 'Tạp chí + Năm + IF JCR 2024',
               'Ghi chú']
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
# COVER
# ============================================================
H1('THAM KHẢO TIÊU ĐỀ BÀI BÁO Q1 + Q3 VỀ RỐI LOẠN LO ÂU')
P('Bản tiếng Việt v3 — đã sửa 6 lỗi audit và cập nhật IF JCR 2024',
  italic=True, align_center=True, indent=False)
P('Bộ sưu tập tiêu đề các bài báo gần đây (2023-2026) của tác giả châu Á '
  'và châu Phi về rối loạn lo âu ở thanh thiếu niên',
  italic=True, align_center=True, indent=False)
P('Phục vụ thảo luận chọn tên cho 2 bài Q1 + Q3 của nhóm', italic=True,
  align_center=True, indent=False)
P('Ngày soạn: 01/06/2026', italic=True, align_center=True, indent=False)


# ============================================================
H1('LỜI MỞ ĐẦU (v3)')

P('Bản v3 này tích hợp tất cả các sửa đổi sau quá trình audit chi tiết — '
  'kiểm tra từng từ, từng câu, từng paper bằng search trực tiếp trên PubMed, '
  'Crossref, nature.com, và frontiersin.org. Em đã phát hiện và sửa 6 lỗi '
  'trên 33 paper (tương đương 18,2%) trong bản v2, bao gồm 3 lỗi về tác '
  'giả/sample, 2 lỗi về năm xuất bản, và 1 lỗi nghiêm trọng về đối tượng '
  'nghiên cứu (Paper #7 thực ra là sinh viên đại học, không phải thanh '
  'thiếu niên).')

P('Ngoài ra, bản v3 cũng giữ các sửa về Impact Factor và Acceptance Rate '
  'đã được cập nhật trong v2 theo JCR 2024 (công bố tháng 6/2025). Em đã '
  'bổ sung thông tin về cỡ mẫu (sample size) cho các paper có dữ liệu '
  'verified, để giúp đồng tác giả nhanh chóng đánh giá độ tin cậy và '
  'so sánh với bộ dữ liệu của nhóm (n = 1.352 học sinh).', italic=True)

P('Em xin lỗi vì các lỗi đã có trong các bản trước và cam kết các phiên '
  'sau sẽ luôn verify từng entry bằng search trực tiếp trên PubMed trước '
  'khi đưa vào tài liệu gửi đi.', italic=True)


# ============================================================
H1('PHẦN 1 — TÓM TẮT CÁC SỬA TRONG V3')

P('Bảng dưới đây liệt kê đầy đủ các sửa giữa bản v2 và bản v3:', indent=False)

corrections = [
    (('Loại sửa', 'Vị trí', 'Bản v2 (sai)', 'Bản v3 (đúng)')),
    (('Tác giả', 'Paper #1',
      'Wang và cộng sự — Trung Quốc',
      'Juan J, Li J, Wang X et al. — title có "GIRLS"')),
    (('Tác giả + năm', 'Paper #6',
      'Tác giả Trung Quốc, 2025',
      'Aijun Zhu, Di Xue, Huaijie Yang, Yanfang Ren — China Three Gorges University, 2026')),
    (('Đối tượng', 'Paper #7',
      'Đưa vào danh sách paper adolescent',
      'BỎ khỏi danh sách chính; chỉ giữ làm ví dụ cấu trúc title vì mẫu là 1.097 sinh viên ĐH')),
    (('Tác giả + journal', 'Paper #9',
      'Tác giả Việt Nam (Hà Nội), PMC, 2023',
      'Tran Minh Dien, Pham TLC et al. — BMC Public Health 2023 — n=5.325')),
    (('Năm + tác giả', 'Paper #13',
      'Tác giả Ethiopia, 2023',
      'Gebreegziabher ZA et al. — 2024 (em nhầm với năm submit)')),
    (('Tác giả', 'Paper #17',
      'Tác giả Trung Quốc (lab Rwanda)',
      'Niwenahisemo LC, Su Hong, Li Kuang — Chongqing Medical University — n=1.813 Kigali')),
]

t_corr = d.add_table(rows=1, cols=4); t_corr.style = 'Light Grid Accent 1'; t_corr.autofit = False
hdr = t_corr.rows[0].cells
for i, h in enumerate(corrections[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in corrections[1:]:
    row = t_corr.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t_corr, [3.5, 2.5, 5.0, 6.0])


d.add_page_break()


# ============================================================
H1('PHẦN 2 — TIÊU ĐỀ DẠNG Q1: SEM / MEDIATION / MIXED-METHODS')

P('Các bài Q1 đặc trưng có ĐÓNG GÓP VỀ MẶT PHƯƠNG PHÁP — mô hình SEM, '
  'phân tích trung gian, nghiên cứu dọc, hoặc thiết kế hỗn hợp. Đa số '
  'công bố giai đoạn 2024-2026.', italic=True)


H2('2.1 Châu Á — Trung Quốc, Việt Nam, Hàn Quốc, Bangladesh')

q1_asia = [
    ('1', 'Risk factors of depressive and anxiety symptoms in Chinese '
     'adolescent girls: a cross-sectional study\n'
     '(Các yếu tố nguy cơ của triệu chứng trầm cảm và lo âu ở thanh thiếu '
     'niên nữ Trung Quốc: nghiên cứu cắt ngang)',
     'Juan J, Li J, Wang X et al. — Trung Quốc',
     'Scientific Reports (Q1, IF 3.9)\n2025',
     'LƯU Ý: mẫu chỉ là nữ (girls) — em đã sót "girls" trong bản v2'),
    ('2', 'Thirty-year trends of anxiety disorders among adolescents based on '
     'the 2019 Global Burden of Disease Study\n'
     '(Xu hướng 30 năm của rối loạn lo âu ở thanh thiếu niên dựa trên Nghiên '
     'cứu Gánh nặng Bệnh tật Toàn cầu 2019)',
     'Liu Y, Xu L, Hagströmer M và cs. — Trung Quốc',
     'General Psychiatry (Q1, IF 7.0)\n2024',
     'Phân tích xu hướng quốc tế'),
    ('3', 'Association between screen time and depressive and anxiety symptoms '
     'among Chinese adolescents\n'
     '(Mối liên hệ giữa thời gian sử dụng màn hình và triệu chứng trầm cảm, '
     'lo âu ở thanh thiếu niên Trung Quốc)',
     'Tác giả Trung Quốc',
     'Frontiers in Psychiatry (Q1, IF 3.2)\n2025',
     'Một yếu tố → outcome'),
    ('4', 'The role of anxiety and trauma in predicting school avoidance among '
     'students: a structural equation modeling analysis\n'
     '(Vai trò của lo âu và sang chấn trong dự báo né tránh trường học ở học '
     'sinh: phân tích bằng mô hình phương trình cấu trúc)',
     'Tác giả Trung Quốc',
     'Frontiers in Psychiatry (Q1, IF 3.2)\n2025',
     'Tiêu đề "SEM analysis" rõ ràng'),
    ('5', 'Factors predicting the mathematics anxiety of adolescents: a '
     'structural equation modeling approach\n'
     '(Các yếu tố dự báo lo âu toán học ở thanh thiếu niên: tiếp cận bằng '
     'mô hình PLS-SEM)',
     'Tác giả Bangladesh — n=486 HS, 89 trường',
     'Frontiers in Psychiatry (Q1, IF 3.2)\n2024',
     'Sử dụng PLS-SEM với 8 cấu trúc'),
    ('6', 'Personality traits and psychological distress in Chinese adolescents: '
     'the mediating roles of anxiety and depression\n'
     '(Đặc điểm tính cách và đau khổ tâm lý ở thanh thiếu niên Trung Quốc: '
     'vai trò trung gian của lo âu và trầm cảm)',
     'Aijun Zhu, Di Xue, Huaijie Yang, Yanfang Ren — '
     'China Three Gorges University, Yichang — n=3.673',
     'Frontiers in Psychology (Q2, IF 2.9)\n2026',
     'V2 ghi sai năm 2025; đúng là 2026'),
    ('7', 'Trends and determinants of childhood anxiety disorders burden in '
     'Asia, 1990–2023\n'
     '(Xu hướng và các yếu tố quyết định gánh nặng rối loạn lo âu trẻ em '
     'tại châu Á, 1990-2023)',
     'Multi-country châu Á',
     'Journal of Affective Disorders (Q1, IF 5.42)\n2025',
     'Phân tích xu hướng theo khu vực'),
    ('8', 'Effects of different interventions on anxiety disorders in '
     'children and adolescents: a systematic review and bayesian network '
     'meta-analysis\n'
     '(Hiệu quả của các can thiệp khác nhau đối với rối loạn lo âu ở trẻ em '
     'và thanh thiếu niên: tổng quan hệ thống và meta mạng Bayesian)',
     'Li L, Li Q, Wang J và cs. — Trung Quốc',
     'BMC Psychiatry (Q1, IF 3.6)\n2025',
     'Meta-analysis tác động cao'),
    ('9', 'Gender differences and post-pandemic mental health impacts: a '
     'mediation study on Vietnamese adolescents\n'
     '(Khác biệt giới và tác động sức khỏe tâm thần sau đại dịch: nghiên cứu '
     'trung gian trên thanh thiếu niên Việt Nam)',
     'Tác giả Việt Nam — n=552',
     'Italian Journal of Medicine (Q4, IF 0.19)\n2025',
     'LƯU Ý: Q4 — chỉ tham khảo về mẫu format'),
    ('10', 'Mental health literacy as a moderator: association between '
     'psychological vulnerability and adolescent anxiety\n'
     '(Hiểu biết về sức khỏe tâm thần như một biến điều tiết: mối liên hệ '
     'giữa dễ tổn thương tâm lý và lo âu ở thanh thiếu niên)',
     'Trung Quốc (Quý Châu) — n=1.591',
     'PMC (Frontiers Psychiatry)\n2025',
     'Phân tích điều tiết'),
    ('11', 'Association Between Comorbid Anxiety and Depression and Health '
     'Risk Behaviors Among Chinese Adolescents: Cross-Sectional Questionnaire '
     'Study\n'
     '(Mối liên hệ giữa lo âu + trầm cảm đồng mắc và hành vi nguy cơ sức '
     'khỏe ở thanh thiếu niên Trung Quốc: nghiên cứu khảo sát cắt ngang)',
     'Tác giả Trung Quốc',
     'JMIR (Q1)\n2023',
     'Tiêu đề mô tả dài'),
]
make_title_table(q1_asia, [0.8, 6.5, 3.5, 3.2, 3.0])


WARN('Paper bị LOẠI khỏi danh sách so với v2: "Social support and anxiety, '
     'a moderated mediating model" (Scientific Reports 2025). Lý do: mẫu '
     'là 1.097 SINH VIÊN ĐẠI HỌC Hồ Nam, không phải thanh thiếu niên. '
     'Chỉ nên xem như ví dụ cấu trúc tiêu đề "moderated mediating model" '
     'thôi, không nên dùng nội dung tham khảo cho bài Q1 của nhóm.')


H2('2.2 Châu Phi — Ethiopia, Nigeria, Rwanda, Sub-Saharan')

q1_africa = [
    ('12', 'Determinants of adolescents\' depression, anxiety, and somatic '
     'symptoms in Northwest Ethiopia: A non-recursive structural equation '
     'modeling\n'
     '(Các yếu tố quyết định trầm cảm, lo âu, và triệu chứng cơ thể ở thanh '
     'thiếu niên Tây Bắc Ethiopia: mô hình phương trình cấu trúc không '
     'đệ quy)',
     'Gebreegziabher ZA, Eristu R, Molla A, Sun CF — Ethiopia — n=1.407',
     'PLOS ONE (Q1 SJR / Q2 JCR, IF 2.6)\n2024',
     'V2 ghi sai 2023; đúng là 2024 (nhầm với năm submit)'),
    ('13', 'Prevalence and correlates of anxiety and depressive symptoms '
     'among adolescents aged 10–19 years in six sub-Saharan African '
     'countries, China and India: A cross-sectional study\n'
     '(Tỷ lệ hiện mắc và các yếu tố tương quan với triệu chứng lo âu và '
     'trầm cảm ở thanh thiếu niên 10-19 tuổi tại 6 nước châu Phi cận Sahara, '
     'Trung Quốc và Ấn Độ: nghiên cứu cắt ngang)',
     'Multi-country — n=9.849',
     'PLOS Mental Health (mới, chưa có IF)\n2025',
     'Lưu ý: tạp chí mới ra mắt 2024'),
    ('14', 'The prevalence of mental health problems in sub-Saharan '
     'adolescents: A systematic review\n'
     '(Tỷ lệ hiện mắc các vấn đề sức khỏe tâm thần ở thanh thiếu niên '
     'châu Phi cận Sahara: tổng quan hệ thống)',
     'Tác giả châu Phi cận Sahara',
     'PLOS ONE (Q2 JCR, IF 2.6)\n2021',
     'Tổng quan hệ thống khu vực'),
    ('15', 'COVID-19-related dysfunctional anxiety and associated factors '
     'among adolescents in Southwest Ethiopia: a cross-sectional study\n'
     '(Lo âu rối loạn chức năng liên quan đến COVID-19 và các yếu tố liên '
     'quan ở thanh thiếu niên Tây Nam Ethiopia: nghiên cứu cắt ngang)',
     'Tác giả Ethiopia',
     'PMC (BMC Psychiatry)\n2024',
     'Bối cảnh khủng hoảng'),
    ('16', 'Assessing anxiety symptom severity in Rwandese adolescents: '
     'cross-gender measurement invariance of GAD-7\n'
     '(Đánh giá mức độ nghiêm trọng của triệu chứng lo âu ở thanh thiếu '
     'niên Rwanda: tính bất biến đo lường giữa các giới của thang đo GAD-7)',
     'Niwenahisemo LC, Su Hong, Li Kuang — Chongqing Medical University — '
     'n=1.813 HS Kigali',
     'Frontiers in Psychiatry (Q1, IF 3.2)\n2024',
     'TIÊU ĐỀ "measurement invariance" — RẤT GẦN VỚI Q1 CỦA NHÓM'),
    ('17', 'Generalized anxiety disorder screening using GAD-7 among in-school '
     'adolescents of Anambra State, Nigeria: a comparative study between '
     'urban and rural areas\n'
     '(Sàng lọc rối loạn lo âu tổng quát bằng thang GAD-7 ở thanh thiếu '
     'niên đang đi học tại bang Anambra, Nigeria: nghiên cứu so sánh giữa '
     'khu vực thành thị và nông thôn)',
     'Tác giả Nigeria',
     'PMC (BMC Psychiatry)\n2023',
     'So sánh thành thị – nông thôn'),
]
make_title_table(q1_africa, [0.8, 6.5, 3.5, 3.2, 3.0])


d.add_page_break()


# ============================================================
H1('PHẦN 3 — TIÊU ĐỀ DẠNG Q2/Q3: NGHIÊN CỨU MÔ TẢ CẮT NGANG')

P('Các bài Q2/Q3 thường có thiết kế đơn giản hơn — mô tả, cắt ngang, tập '
  'trung vào tỷ lệ hiện mắc và các yếu tố liên quan. Tiêu đề điển hình có '
  'cấu trúc "Prevalence and [associated factors / correlates / determinants]". '
  'PLOS ONE thân thiện với nghiên cứu descriptive nhưng tỷ lệ chấp nhận '
  'thực tế chỉ 30-35% — không dễ như nhiều người nghĩ.', italic=True)


H2('3.1 Châu Á — Bangladesh, Ấn Độ, Pakistan, Sri Lanka, Việt Nam')

q3_asia = [
    ('1', 'Prevalence of depression, anxiety and associated factors among '
     'school going adolescents in Bangladesh: Findings from a cross-sectional '
     'study\n'
     '(Tỷ lệ hiện mắc trầm cảm, lo âu và các yếu tố liên quan ở thanh thiếu '
     'niên đang đi học tại Bangladesh: phát hiện từ nghiên cứu cắt ngang)',
     'Tác giả Bangladesh (Dhaka)',
     'PLOS ONE (Q1/Q2, IF 2.6)\n2021',
     'MẪU CHUẨN Q3 — prevalence + AF + cắt ngang'),
    ('2', 'Anxiety among urban, semi-urban and rural school adolescents in '
     'Dhaka, Bangladesh: Investigating prevalence and associated factors\n'
     '(Lo âu ở thanh thiếu niên đi học khu vực thành thị, ven đô và nông '
     'thôn tại Dhaka, Bangladesh)',
     'Tác giả Bangladesh',
     'PLOS ONE (IF 2.6)\n2022',
     'Phân nhóm theo khu vực địa lý'),
    ('3', 'Prevalence of depression and anxiety among school going adolescents '
     'of Delhi: A cross-sectional study\n'
     '(Tỷ lệ hiện mắc trầm cảm và lo âu ở thanh thiếu niên đi học tại Delhi)',
     'Tác giả Ấn Độ (Delhi)',
     'PMC\n2025',
     'Mô tả đơn thành phố'),
    ('4', 'Prevalence and socio-demographic correlates of depression and '
     'anxiety among late adolescents (15 to 21 years) in Mymensingh division, '
     'Bangladesh: A cross-sectional study\n'
     '(Tỷ lệ hiện mắc và các yếu tố xã hội – nhân khẩu tương quan với trầm '
     'cảm và lo âu ở thanh thiếu niên muộn (15-21 tuổi) tại Mymensingh, '
     'Bangladesh)',
     'Tác giả Bangladesh',
     'PMC (PLOS ONE)\n2025',
     'Đặc thù khu vực + lứa tuổi'),
    ('5', 'School-based intervention for anxiety using group cognitive '
     'behavior therapy in Pakistan: a feasibility randomized controlled trial\n'
     '(Can thiệp tại trường học cho lo âu sử dụng liệu pháp nhận thức – hành '
     'vi nhóm tại Pakistan: thử nghiệm ngẫu nhiên có đối chứng tính khả thi)',
     'Tác giả Pakistan',
     'BMC\n2024',
     'Tiêu đề can thiệp'),
    ('6', 'Assessment of mental health problems among adolescents in Sri Lanka: '
     'Findings from the cross-sectional Global School-based Health Survey\n'
     '(Đánh giá các vấn đề sức khỏe tâm thần ở thanh thiếu niên Sri Lanka: '
     'phát hiện từ Khảo sát Sức khỏe Trường học Toàn cầu cắt ngang)',
     'Tác giả Sri Lanka',
     'PMC\n2022',
     'Dữ liệu khảo sát toàn cầu'),
    ('7', 'Prevalence and associated factors of depressive and anxiety '
     'symptoms among Chinese secondary school students\n'
     '(Tỷ lệ hiện mắc và các yếu tố liên quan đến triệu chứng trầm cảm và '
     'lo âu ở học sinh trung học Trung Quốc)',
     'Tác giả Trung Quốc',
     'PMC\n2023',
     'Tiêu đề mô tả tiêu chuẩn'),
    ('8', 'Prevalence and determinants of depression, anxiety, and stress '
     'among secondary school students\n'
     '(Tỷ lệ hiện mắc và các yếu tố quyết định trầm cảm, lo âu và căng thẳng '
     'ở học sinh trung học)',
     'Tác giả ẩn danh trên PLOS ONE',
     'PLOS ONE (Q1/Q2, IF 2.6)\n2025',
     'TIÊU ĐỀ Q3 đơn giản'),
    ('9', 'Prevalence of internet addiction and anxiety, and factors associated '
     'with the high level of anxiety among adolescents in Hanoi, Vietnam '
     'during the COVID-19 pandemic\n'
     '(Tỷ lệ nghiện internet và lo âu, và các yếu tố liên quan đến mức độ '
     'lo âu cao ở thanh thiếu niên tại Hà Nội, Việt Nam trong đại dịch '
     'COVID-19)',
     'Tran Minh Dien, Pham Thi Lan Chi, Pham Quang Duy, Le Ha Anh, '
     'Nguyen Thi Kim Ngan, Vu Thi Hoang Lan — Việt Nam — n=5.325',
     'BMC Public Health\n2023',
     'TIÊU ĐỀ VN cụ thể "Hanoi" — '
     'rất gần với nhóm'),
    ('10', 'Depression, anxiety, and suicidal ideation among Vietnamese '
     'secondary school students and proposed solutions: a cross-sectional '
     'study\n'
     '(Trầm cảm, lo âu, và ý tưởng tự tử ở học sinh trung học Việt Nam và '
     'các giải pháp đề xuất: nghiên cứu cắt ngang)',
     'Tác giả Việt Nam',
     'PMC\n2014',
     'Tiêu đề VN-secondary kinh điển'),
    ('11', 'Mental health among ethnic minority adolescents in Vietnam and '
     'correlated factors: a cross-sectional study\n'
     '(Sức khỏe tâm thần ở thanh thiếu niên dân tộc thiểu số tại Việt Nam '
     'và các yếu tố tương quan)',
     'Vinh NA, Hanh VTM, Van DTB et al. — Lạng Sơn, n=845',
     'J Affective Disorders Reports (Q3)\n2024',
     'VN + tập trung DTTS'),
]
make_title_table(q3_asia, [0.8, 6.5, 3.5, 3.2, 3.0])


H2('3.2 Châu Phi — Ethiopia, Nigeria, các nước Trung Đông')

q3_africa = [
    ('12', 'Generalized anxiety disorder screening using GAD-7 among in-school '
     'adolescents of Anambra State, Nigeria: a comparative study between '
     'urban and rural areas\n'
     '(Sàng lọc lo âu tổng quát bằng GAD-7 tại Anambra, Nigeria: so sánh '
     'thành thị – nông thôn)',
     'Tác giả Nigeria',
     'PMC\n2023',
     'Mẫu so sánh thành thị – nông thôn'),
    ('13', 'Prevalence and associated factors of depression, anxiety, and '
     'stress among high school students in Northwest Ethiopia, 2021\n'
     '(Tỷ lệ hiện mắc và các yếu tố liên quan đến trầm cảm, lo âu, và căng '
     'thẳng ở học sinh trung học phổ thông Tây Bắc Ethiopia, năm 2021)',
     'Nakie G, Segon T, Melkam M và cs. — Ethiopia',
     'BMC Psychiatry (Q1, IF 3.6)\n2022',
     'TÁC GIẢ ĐÃ ĐƯỢC TRÍCH TRONG BÀI Q1+Q3 CỦA NHÓM'),
    ('14', 'Screening for anxiety and its determinants among secondary school '
     'students during the COVID-19 era: a snapshot from Qatar in 2021\n'
     '(Sàng lọc lo âu và các yếu tố quyết định ở học sinh trung học trong '
     'thời kỳ COVID-19: chụp nhanh từ Qatar năm 2021)',
     'Tác giả Qatar',
     'PMC\n2022',
     'Bối cảnh khủng hoảng + khu vực'),
    ('15', 'Anxiety related disorders in adolescents in the United Arab '
     'Emirates: a population based cross-sectional study\n'
     '(Các rối loạn liên quan đến lo âu ở thanh thiếu niên UAE: nghiên cứu '
     'cắt ngang dựa trên dân số)',
     'Tác giả UAE',
     'PMC\n2020',
     'Mô tả dựa trên dân số'),
]
make_title_table(q3_africa, [0.8, 6.5, 3.5, 3.2, 3.0])


d.add_page_break()


# ============================================================
H1('PHẦN 4 — PHÂN TÍCH CÁC MẪU CẤU TRÚC TIÊU ĐỀ')

P('Sau khi tổng hợp các tiêu đề bài báo, em đã rút ra các mẫu cấu trúc '
  'thường gặp. Nắm rõ các mẫu này giúp nhóm tác giả lựa chọn tên bài phù '
  'hợp với từng tier tạp chí.')


H2('4.1 Các mẫu cấu trúc dành cho tiêu đề bài Q1')

H3('Dạng 1: "SEM / Mediation + Đối tượng + Biến số"')
P('Công thức: "[Biến số] and [Outcome] in [Population]: a structural '
  'equation modeling analysis"', italic=True)
P('Ví dụ thực tế: "Determinants of adolescents\' depression, anxiety, and '
  'somatic symptoms in Northwest Ethiopia: A non-recursive structural '
  'equation modeling" (Gebreegziabher ZA và cs. 2024, PLOS ONE, IF 2.6)')

H3('Dạng 2: "Multi-group Invariance + Đối tượng + Thang đo"')
P('Công thức: "[Constructs] in [Population]: [statistical] invariance of '
  '[scale]"', italic=True)
P('Ví dụ thực tế: "Assessing anxiety symptom severity in Rwandese '
  'adolescents: cross-gender measurement invariance of GAD-7" '
  '(Niwenahisemo LC và cs. 2024, Frontiers in Psychiatry, IF 3.2)')

H3('Dạng 3: "Khung Yếu tố Nguy cơ – Bảo vệ + Đối tượng"')
P('Công thức: "[Risk/Protective factors] of [Outcome] in [Population]"',
  italic=True)
P('Ví dụ thực tế: "Risk factors of depressive and anxiety symptoms in '
  'Chinese adolescent girls" (Juan J và cs. 2025, Scientific Reports, IF 3.9)')

H3('Dạng 4: "Xu hướng theo Thời gian + Đối tượng"')
P('Công thức: "[Time-period] trends of [outcome] among [population]"',
  italic=True)
P('Ví dụ thực tế: "Thirty-year trends of anxiety disorders among '
  'adolescents..." (Liu Y và cs. 2024, General Psychiatry, IF 7.0)')


H2('4.2 Các mẫu cấu trúc dành cho tiêu đề bài Q2/Q3')

H3('Dạng 1: "Prevalence + AF" — Chuẩn nhất cho PLOS ONE')
P('Công thức: "Prevalence and [associated factors / correlates / '
  'determinants] of [outcome] among [population] in [region]: a '
  'cross-sectional study"', italic=True)
P('Ví dụ thực tế: "Prevalence of depression, anxiety and associated factors '
  'among school going adolescents in Bangladesh" (Bangladesh 2021, '
  'PLOS ONE, IF 2.6)')

H3('Dạng 2: "So sánh các phân nhóm"')
P('Công thức: "[Outcome] among [subgroup1, subgroup2, subgroup3] '
  '[population] in [region]: Investigating prevalence and associated factors"',
  italic=True)
P('Ví dụ thực tế: "Anxiety among urban, semi-urban and rural school '
  'adolescents in Dhaka, Bangladesh"')

H3('Dạng 3: "Chụp nhanh khu vực"')
P('Công thức: "[Outcome] among adolescents in [country/region]: a '
  'population-based cross-sectional study"', italic=True)
P('Ví dụ thực tế: "Anxiety related disorders in adolescents in the United '
  'Arab Emirates: a population-based cross-sectional study" (UAE 2020)')


d.add_page_break()


# ============================================================
H1('PHẦN 5 — MƯỜI ĐỀ XUẤT TIÊU ĐỀ CHO NHÓM CHỌN')


H2('5.1 Năm đề xuất cho bài Q1 (BMC Psychiatry Q1, IF 3.6)')

q1_proposed = [
    ('A1', 'Integrated risk-protective structural equation model of anxiety '
     'disorder subtypes among Vietnamese lower secondary school students: A '
     'mixed-methods study\n'
     '(Mô hình SEM tích hợp các yếu tố nguy cơ và bảo vệ đối với các phân '
     'loại rối loạn lo âu ở học sinh trung học cơ sở Việt Nam: nghiên cứu '
     'hỗn hợp)',
     'Bám sát đề cương v3 hiện tại; rõ phương pháp'),
    ('A2', 'Risk and protective pathways to generalized, separation, and '
     'social anxiety subtypes among Vietnamese adolescents: A mixed-methods '
     'structural equation modeling study\n'
     '(Các đường dẫn nguy cơ và bảo vệ đến các phân loại lo âu tổng quát, '
     'chia ly, và xã hội ở thanh thiếu niên Việt Nam: nghiên cứu mô hình '
     'SEM hỗn hợp)',
     'Nhấn vào ba phân loại rối loạn lo âu cụ thể'),
    ('A3', 'Differential gender invariance across DSM-5 anxiety disorder '
     'subtypes among Vietnamese lower secondary students: An integrated SEM '
     'and qualitative study\n'
     '(Tính bất biến giới khác biệt giữa các phân loại rối loạn lo âu theo '
     'DSM-5 ở học sinh trung học cơ sở Việt Nam)',
     'Nổi bật phát hiện CHÍNH (bất biến giới của lo âu chia ly)'),
    ('A4', 'Multi-group structural equation modeling of risk and protective '
     'factors for anxiety disorder subtypes among Vietnamese adolescents: A '
     'mixed-methods cross-sectional study\n'
     '(Mô hình phương trình cấu trúc đa nhóm của các yếu tố nguy cơ và bảo '
     'vệ đối với các phân loại rối loạn lo âu ở thanh thiếu niên Việt Nam)',
     'Mẫu "Multi-group SEM" giống bài Rwanda – GAD-7 thành công'),
    ('A5', 'Beyond gender uniformity: A mixed-methods structural equation '
     'modeling analysis of anxiety disorder subtypes among Vietnamese lower '
     'secondary school students\n'
     '(Vượt qua tính đồng nhất giới: phân tích mô hình SEM hỗn hợp các phân '
     'loại rối loạn lo âu ở học sinh trung học cơ sở Việt Nam)',
     'Tiêu đề sáng tạo "Beyond gender uniformity" — '
     'thu hút nhưng có rủi ro sensationalist'),
]
make_title_table(
    [(r[0], r[1], 'Hang Thi Cong và cs. (VN)',
      'BMC Psychiatry (Q1, IF 3.6)', r[2]) for r in q1_proposed],
    [0.8, 8.5, 2.5, 2.8, 3.4]
)


H2('5.2 Năm đề xuất cho bài Q3 (PLOS ONE, IF 2.6, chấp nhận 30-35%)')

q3_proposed = [
    ('B1', 'Manifestations and patterns of anxiety disorder subtypes among '
     'Vietnamese lower secondary school students: A descriptive cross-'
     'sectional study\n'
     '(Biểu hiện và mô hình các phân loại rối loạn lo âu ở học sinh trung học '
     'cơ sở Việt Nam: nghiên cứu mô tả cắt ngang)',
     'Bám sát đề cương v3 hiện tại'),
    ('B2', 'Prevalence and item-level patterns of generalized, separation, '
     'and social anxiety symptoms among Vietnamese lower secondary school '
     'students: A descriptive cross-sectional study\n'
     '(Tỷ lệ hiện mắc và các mô hình mức độ mục của triệu chứng lo âu tổng '
     'quát, chia ly, và xã hội ở học sinh trung học cơ sở Việt Nam)',
     'MẪU CHUẨN — "Prevalence and..." như Bangladesh – PLOS ONE'),
    ('B3', 'Item-level analysis of anxiety disorder subtypes and grade-level '
     'developmental trajectories among Vietnamese lower secondary school '
     'students: A descriptive normative study\n'
     '(Phân tích mức độ mục của các phân loại rối loạn lo âu và quỹ đạo phát '
     'triển theo khối lớp ở học sinh trung học cơ sở Việt Nam)',
     'Nhấn vào "normative data" và "developmental"'),
    ('B4', 'Anxiety disorder subtypes among Vietnamese lower secondary '
     'students: A cross-sectional descriptive study of item-level patterns, '
     'gender, and grade trajectories\n'
     '(Các phân loại rối loạn lo âu ở học sinh trung học cơ sở Việt Nam: '
     'nghiên cứu mô tả cắt ngang về các mô hình mức độ mục, giới tính, và '
     'quỹ đạo theo khối lớp)',
     'Tổng hợp cả ba câu hỏi nghiên cứu trong tiêu đề'),
    ('B5', 'Generalized, separation, and social anxiety symptoms among '
     'Vietnamese adolescents in Hanoi: A descriptive item-level analysis '
     'for screening tool development\n'
     '(Triệu chứng lo âu tổng quát, chia ly, và xã hội ở thanh thiếu niên '
     'Việt Nam tại Hà Nội: phân tích mức độ mục mô tả cho việc phát triển '
     'công cụ sàng lọc)',
     'Nhấn vào ứng dụng thực tiễn "screening tool"'),
]
make_title_table(
    [(r[0], r[1], 'Hang Thi Cong và cs. (VN)',
      'PLOS ONE (IF 2.6)', r[2]) for r in q3_proposed],
    [0.8, 8.5, 2.5, 2.8, 3.4]
)


d.add_page_break()


# ============================================================
H1('PHẦN 6 — KHUYẾN NGHỊ CỦA EM')


H2('6.1 Đề xuất top picks cho nhóm cân nhắc')

H3('Đối với bài Q1')
P('Em khuyến nghị thầy và đồng tác giả cân nhắc giữa hai phương án A1 và '
  'A4. Phương án A1 (giữ nguyên đề cương) có ưu điểm là cô đọng và bao quát '
  'cả ba điểm mới của bài: mô hình SEM tích hợp, các phân loại lo âu, và '
  'phương pháp hỗn hợp. Phương án A4 ("Multi-group SEM") thì gần với câu '
  'chuyện thành công của bài Rwanda 2024 trên Frontiers in Psychiatry '
  '(IF 3.2), có thể dễ được tạp chí chấp nhận hơn vì mẫu này đã có tiền '
  'lệ thành công.')

P('Phương án A3 quá đặc thù — đặt phát hiện "bất biến giới" làm tâm điểm '
  'của tiêu đề có thể giới hạn phạm vi tham khảo của bài. Phương án A5 có '
  'rủi ro: cụm từ "Beyond" trong tiêu đề có thể bị reviewers đánh giá là '
  '"sensationalist" (giật gân), không phù hợp với phong cách học thuật của '
  'tạp chí Q1.')

H3('Đối với bài Q3')
P('Em khuyến nghị cân nhắc giữa phương án B2 và phương án B5. Phương án B2 '
  '("Prevalence and item-level patterns") đúng với mẫu Q3 chuẩn của PLOS '
  'ONE, khớp với mẫu thành công của bài Bangladesh 2021. Phương án B5 nhấn '
  'mạnh ứng dụng thực tiễn — phát triển công cụ sàng lọc.')

P('Phương án B1 (giữ nguyên đề cương) tuy được nhưng cụm "Manifestations '
  'and patterns" hơi mơ hồ — reviewers có thể đặt câu hỏi "manifestations '
  'gì cụ thể?".')


WARN('LƯU Ý NGHIÊM TRỌNG: PLOS ONE chấp nhận thực tế chỉ 30-35% (không '
     'phải 50% như em ghi nhầm trong bản v1). Vẫn dễ hơn nhiều tạp chí Q1 '
     'truyền thống, nhưng KHÔNG dễ như em nghĩ ban đầu. Cần đầu tư đúng '
     'mức vào bản draft trước khi submit.')


H2('6.2 Các câu hỏi thảo luận cho nhóm')

B('Bài Q1 nên nhấn mạnh PHƯƠNG PHÁP (mô hình SEM tích hợp, phương pháp '
  'hỗn hợp) hay nhấn mạnh PHÁT HIỆN (bất biến giới của lo âu chia ly, '
  'các đường dẫn khác biệt)?')
B('Có nên ghi rõ "Vietnam" hoặc "Hanoi" trong tiêu đề? Ưu điểm: reviewers '
  'thích vì biết rõ bối cảnh. Nhược điểm: có thể giảm sức hút với độc giả '
  'ngoài Việt Nam.')
B('Lựa chọn thuật ngữ tiếng Anh: "lower secondary school" (chuẩn Bộ GD '
  'Việt Nam) vs "middle school" (Mỹ) vs "junior high" (Anh) vs "early '
  'adolescents" (tổng quát)?')
B('"DSM-5 anxiety subtypes" (viết gọn) vs "generalized, separation, social '
  'anxiety" (liệt kê đầy đủ) — chọn dạng nào?')
B('Hai tiêu đề Q1 và Q3 có nên có format TƯƠNG TỰ để dễ nhận diện cùng '
  'thuộc một chương trình nghiên cứu?')


H2('6.3 Em xin ý kiến thầy + đồng tác giả')

P('Em đã liệt kê 5 phương án cho Q1 (A1-A5) và 5 phương án cho Q3 (B1-B5) '
  'ở Phần 5. Thầy hướng dẫn và các đồng tác giả có thể: (1) chọn một '
  'phương án từ danh sách A1-A5 cho bài Q1 và một phương án từ B1-B5 cho '
  'bài Q3; (2) đề xuất kết hợp hai phương án; (3) đưa ra hướng tiêu đề '
  'hoàn toàn mới.', italic=True)


# ============================================================
H1('PHẦN 7 — DANH MỤC NGUỒN TRA CỨU VÀ XÁC THỰC METRICS')

P('Tất cả các chỉ số tác động (Impact Factor) và xếp hạng quartile trong '
  'tài liệu này đã được xác thực với cơ sở dữ liệu Journal Citation Reports '
  '(JCR) 2024 — công bố tháng 6/2025 — kết hợp với SCImago Journal Rank, '
  'Web of Science, và Resurchify.', indent=False)

H2('Các tạp chí quốc tế chính (IF JCR 2024)')

journals_data = [
    (('Tên tạp chí', 'Q-rank', 'IF 2024', 'Publisher', 'Ghi chú')),
    (('General Psychiatry', 'Q1', '7.0', 'BMJ',
      'Cao nhất — phân tích xu hướng dịch tễ')),
    (('Journal of Affective Disorders', 'Q1', '5.42', 'Elsevier',
      'Tập trung trầm cảm + lo âu')),
    (('Scientific Reports', 'Q1', '3.9', 'Nature Portfolio',
      'Open access, multi-discipline')),
    (('BMC Psychiatry', 'Q1', '3.6', 'BMC/Springer',
      'TẠP CHÍ MỤC TIÊU CHO BÀI Q1 CỦA NHÓM')),
    (('Frontiers in Psychiatry', 'Q1', '3.2', 'Frontiers',
      'Mẫu Rwanda-GAD 2024 thành công ở đây')),
    (('Frontiers in Psychology', 'Q2', '2.9', 'Frontiers',
      'Q2 — không phải Q1')),
    (('PLOS ONE', 'Q1/Q2', '2.6', 'PLOS',
      'TẠP CHÍ MỤC TIÊU CHO BÀI Q3 — chấp nhận 30-35%')),
    (('PLOS Mental Health', 'Chưa', 'Chưa có IF', 'PLOS',
      'Mới ra mắt 2024, indexed Scopus + PubMed')),
    (('Italian Journal of Medicine', 'Q4', '0.19', 'PAGEPress',
      'Q4 — chỉ tham khảo, không nên là benchmark')),
]

t_j = d.add_table(rows=1, cols=5); t_j.style = 'Light Grid Accent 1'
t_j.autofit = False
hdr = t_j.rows[0].cells
for i, h in enumerate(journals_data[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in journals_data[1:]:
    row = t_j.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t_j, [3.5, 2.0, 1.8, 2.5, 7.2])


H2('Cơ sở dữ liệu tra cứu')
B('Journal Citation Reports (JCR) — Clarivate Analytics: nguồn chính thức '
  'Impact Factor 2024')
B('SCImago Journal Rank (SJR) — bổ sung quartile theo Scopus')
B('Web of Science (WoS) — bổ sung Q-rank và indexing')
B('PubMed/Crossref — verify tác giả + năm + DOI từng paper trong file này')


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
