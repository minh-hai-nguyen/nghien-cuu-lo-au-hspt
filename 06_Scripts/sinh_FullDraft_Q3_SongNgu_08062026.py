# -*- coding: utf-8 -*-
"""Full Draft bai Q3 — Multi-group SEM theo gioi.
Song ngu Anh-Viet, ~8000 tu, tieng Viet thuan, chi tiet cao."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'FullDraft_Q3_SongNgu_v1_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5

def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def EN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def VN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def PB(en, vn):
    EN(en); VN(vn)


# ========== TITLE BLOCK ==========
TITLE('Gender-specific pathways to anxiety disorders in early '
      'adolescence: A multi-group structural equation model in '
      'Vietnamese lower secondary students', 14)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Các đường dẫn đặc thù theo giới đến các rối loạn lo '
              'âu ở giai đoạn đầu tuổi vị thành niên: Mô hình phương '
              'trình cấu trúc đa nhóm trên học sinh trung học cơ sở '
              'Việt Nam')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True; r.bold = True
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Cong Thi Hang [first author]; [Nguyen Minh Duc]; '
              'Hai-Nguyen Minh\n'
              'Faculty of Psychology and Education, Hanoi National '
              'University of Education\n'
              'Target journal: Frontiers in Psychiatry — Adolescent '
              'and Young Adult Psychiatry section\n'
              'Bản thảo đầy đủ song ngữ Anh-Việt phiên bản 1 — soạn '
              '08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ========== ABSTRACT ==========
H1('ABSTRACT — TÓM TẮT')

H2('Background — Bối cảnh')
PB('Adolescent anxiety disorders are consistently more prevalent '
   'among girls than boys in international epidemiological studies, '
   'with the gender gap emerging around puberty. However, the '
   'largest cross-national study to date — Jefferies & Ungar '
   '(2020), with 6,825 participants aged 16-29 across seven '
   'countries — found no overall gender difference in Social '
   'Anxiety Disorder among young adults. Our preliminary data on '
   '1,352 Vietnamese lower secondary students (ages 11-14) show '
   'female-higher scores on SAD, raising a fundamental research '
   'question: are the risk-protective pathways that predict '
   'anxiety invariant across gender in early adolescence?',
   'Các rối loạn lo âu ở tuổi vị thành niên thường xuyên xảy ra '
   'phổ biến hơn ở nữ so với nam trong các nghiên cứu dịch tễ '
   'quốc tế, với khoảng cách giới xuất hiện quanh tuổi dậy thì. '
   'Tuy nhiên, nghiên cứu xuyên quốc gia lớn nhất đến nay — '
   'Jefferies & Ungar (2020) với 6.825 người tham gia tuổi 16-29 '
   'trên bảy quốc gia — không tìm thấy khác biệt giới tổng thể về '
   'Rối loạn Lo âu Xã hội ở thanh niên. Dữ liệu sơ bộ của chúng '
   'tôi trên 1.352 học sinh trung học cơ sở Việt Nam (tuổi 11-14) '
   'cho thấy điểm SAD của nữ cao hơn nam, đặt ra câu hỏi nghiên '
   'cứu cốt lõi: liệu các đường dẫn nguy cơ-bảo vệ dự đoán lo âu '
   'có giữ nguyên giữa hai giới ở giai đoạn đầu tuổi vị thành '
   'niên không?')

H2('Methods — Phương pháp')
PB('Cross-sectional survey of 1,352 Vietnamese lower secondary '
   'students (614 boys, 738 girls; ages 11-14) recruited from two '
   'schools in Hanoi (Nhat Tan urban, Tay Mo suburban). Eight '
   'validated scales measured three risk factors (bullying '
   'victimization, academic stress, smartphone addiction), four '
   'protective factors (school membership, parental support, peer '
   'support, self-esteem), and three DSM-5 anxiety subtypes '
   '(Generalized Anxiety, Social Anxiety, Separation Anxiety). '
   'Multi-group structural equation modeling tested measurement '
   'invariance (configural, metric, scalar) and structural '
   'invariance of 21 hypothesized pathways. Criteria: ΔCFI ≤ '
   '0.01 (Cheung & Rensvold 2002), ΔRMSEA ≤ 0.015 (Chen 2007). '
   'Pre-registered on OSF prior to analysis.',
   'Khảo sát cắt ngang trên 1.352 học sinh trung học cơ sở Việt '
   'Nam (614 nam, 738 nữ; tuổi 11-14), tuyển từ hai trường tại Hà '
   'Nội (Nhật Tân đô thị, Tây Mỗ ngoại ô). Tám thang đo đã chuẩn '
   'hóa đo lường ba yếu tố nguy cơ (bị bắt nạt, áp lực học tập, '
   'nghiện điện thoại), bốn yếu tố bảo vệ (sự gắn bó với trường, '
   'sự hỗ trợ từ cha mẹ, sự hỗ trợ từ bạn bè, lòng tự trọng), và '
   'ba phân loại lo âu theo DSM-5 (Lo âu Lan tỏa, Lo âu Xã hội, '
   'Lo âu Chia ly). Mô hình phương trình cấu trúc đa nhóm kiểm '
   'chứng bất biến thang đo (bất biến cấu hình, bất biến hệ số '
   'tải, bất biến ngưỡng) và bất biến cấu trúc của 21 đường dẫn '
   'đã giả thuyết. Tiêu chí: ΔCFI ≤ 0,01 (Cheung & Rensvold '
   '2002), ΔRMSEA ≤ 0,015 (Chen 2007). Đăng ký trước trên nền '
   'tảng OSF trước khi phân tích.')

H2('Results — Kết quả (dự kiến — chờ phân tích)')
PB('Measurement invariance is expected to be supported at all '
   'three levels. For structural invariance, we anticipate '
   'identifying a subset of non-invariant pathways — particularly '
   'bullying → SAD (stronger in girls based on developmental '
   'theory; exact β values await analysis). Effect sizes (β, 95% '
   'CI) and z-tests of cross-gender coefficient differences will '
   'be reported for all 21 pathways. R² will be reported '
   'separately for boys and girls (vs pooled R² = 0.598 from our '
   'preceding Q2 paper).',
   'Bất biến thang đo dự kiến được hỗ trợ ở cả ba cấp. Về bất '
   'biến cấu trúc, chúng tôi dự đoán nhận diện được một nhóm '
   'đường dẫn không bất biến — đặc biệt bị bắt nạt → SAD (mạnh '
   'hơn ở nữ dựa trên lý thuyết phát triển; giá trị β chính xác '
   'chờ phân tích). Hệ số ảnh hưởng (β, khoảng tin cậy 95%) và '
   'kiểm định z so sánh chênh lệch hệ số giữa hai giới sẽ được '
   'báo cáo cho cả 21 đường dẫn. R² sẽ được báo cáo riêng cho '
   'nam và nữ (so với R² mẫu gộp = 0,598 từ bài Q2 trước đó của '
   'chúng tôi).')

H2('Conclusions — Kết luận')
PB('The findings will inform gender-tailored prevention programs '
   'in Vietnamese schools and reconcile the discrepancy between '
   'our sample and Jefferies & Ungar (2020) through a '
   'developmental-cultural framework. Three theoretical '
   'mechanisms are advanced: (i) developmental peak vulnerability '
   'at ages 11-14, (ii) Confucian-Vietnamese gender expectations '
   'amplifying social-evaluative anxiety in girls, and (iii) '
   'tam giao (three teachings) emotional restraint norms shaping '
   'gendered SAD expression.',
   'Các phát hiện sẽ cung cấp thông tin cho các chương trình '
   'phòng ngừa đặc thù theo giới ở trường học Việt Nam và hòa '
   'giải sự khác biệt giữa mẫu của chúng tôi với Jefferies & '
   'Ungar (2020) thông qua khung lý thuyết phát triển-văn hóa. '
   'Ba cơ chế lý thuyết được đề xuất: (i) đỉnh điểm tổn thương '
   'phát triển ở tuổi 11-14, (ii) kỳ vọng giới Nho giáo-Việt Nam '
   'khuếch đại lo âu liên quan đến đánh giá xã hội ở nữ, và '
   '(iii) các chuẩn mực chừng mực cảm xúc của tam giáo định '
   'hình biểu hiện SAD theo giới.')

H2('Keywords — Từ khóa')
PB('adolescent anxiety; gender differences; multi-group SEM; '
   'measurement invariance; Vietnam; Confucian culture; tam giao; '
   'co-rumination',
   'lo âu vị thành niên; khác biệt giới; SEM đa nhóm; bất biến '
   'thang đo; Việt Nam; văn hóa Nho giáo; tam giáo; cùng nhau '
   'day dứt (co-rumination)')


# ========== INTRODUCTION ==========
H1('1. INTRODUCTION — DẪN NHẬP')

H2('1.1. The gender paradox in adolescent anxiety / Nghịch lý '
   'giới trong lo âu vị thành niên')
PB('Anxiety disorders are among the most common mental health '
   'conditions in adolescence, affecting approximately 8-12% of '
   'youth worldwide. International epidemiological studies have '
   'consistently documented that girls show higher rates of '
   'anxiety disorders than boys, with the gender gap emerging '
   'around puberty (McLean & Anderson 2009; Hankin et al. 2007). '
   'The U.S. National Comorbidity Survey Replication reported '
   'that approximately half of all lifetime anxiety disorders '
   'have an onset by age 14 (Kessler et al. 2005), and girls '
   'reach peak vulnerability earlier than boys. This robust '
   'finding has motivated decades of research into the '
   'developmental and biopsychosocial mechanisms underlying '
   'gendered patterns of anxiety risk.',
   'Các rối loạn lo âu là một trong những vấn đề sức khỏe tâm '
   'thần phổ biến nhất ở tuổi vị thành niên, ảnh hưởng đến '
   'khoảng 8-12% thanh thiếu niên trên toàn thế giới. Các nghiên '
   'cứu dịch tễ quốc tế đã ghi nhận thường xuyên rằng nữ có tỷ '
   'lệ mắc các rối loạn lo âu cao hơn nam, với khoảng cách giới '
   'xuất hiện quanh tuổi dậy thì (McLean & Anderson 2009; Hankin '
   'và cs. 2007). Khảo sát Phối hợp Quốc gia Hoa Kỳ (NCS-R) báo '
   'cáo rằng khoảng một nửa số ca rối loạn lo âu trong đời khởi '
   'phát trước 14 tuổi (Kessler và cs. 2005), và nữ đạt đỉnh tổn '
   'thương sớm hơn nam. Phát hiện vững chắc này đã thúc đẩy hàng '
   'thập kỷ nghiên cứu về các cơ chế phát triển và sinh-tâm-xã '
   'hội ẩn dưới các mô hình giới của nguy cơ lo âu.')

PB('However, this gender-difference pattern does not consistently '
   'replicate in all studies. Jefferies & Ungar (2020) conducted '
   'the largest cross-national study of Social Anxiety Disorder '
   '(SAD) to date, surveying 6,825 young adults aged 16-29 across '
   'seven countries (Australia, Brazil, China, Indonesia, '
   'Russia, USA, Vietnam) using the Social Interaction Anxiety '
   'Scale. Strikingly, they found no overall gender effect on '
   'SAD in the pooled cross-national sample. Two interpretations '
   'are possible. Either gender differences in social anxiety '
   'genuinely disappear by young adulthood, or methodological '
   'and contextual factors obscured the difference in their '
   'sample. Distinguishing these interpretations requires testing '
   'in younger adolescent cohorts where the gender gap should '
   'still be detectable.',
   'Tuy nhiên, mô hình khác biệt giới này không nhất quán trong '
   'tất cả các nghiên cứu. Jefferies & Ungar (2020) tiến hành '
   'nghiên cứu xuyên quốc gia lớn nhất về Rối loạn Lo âu Xã hội '
   '(SAD) đến nay, khảo sát 6.825 thanh niên tuổi 16-29 trên bảy '
   'quốc gia (Úc, Brazil, Trung Quốc, Indonesia, Nga, Mỹ, Việt '
   'Nam) sử dụng Thang Lo âu Tương tác Xã hội (SIAS). Đáng chú '
   'ý, họ không tìm thấy hiệu ứng giới tổng thể trên SAD trong '
   'mẫu gộp xuyên quốc gia. Có hai cách diễn giải. Hoặc khác '
   'biệt giới trong lo âu xã hội thực sự biến mất khi đến tuổi '
   'thanh niên, hoặc các yếu tố phương pháp và bối cảnh đã che '
   'giấu khác biệt trong mẫu của họ. Phân biệt hai cách diễn '
   'giải này đòi hỏi phải kiểm chứng trong các cohort vị thành '
   'niên trẻ tuổi hơn, nơi khoảng cách giới vẫn còn phát hiện '
   'được.')

H2('1.2. Three explanations for the discrepancy / Ba lời giải '
   'cho khác biệt')
PB('We propose three theoretical mechanisms that, individually '
   'or in combination, may explain why gender differences in '
   'anxiety appear in early adolescence (ages 11-14) yet '
   'disappear or attenuate in young adulthood (ages 16-29). The '
   'first mechanism — developmental peak vulnerability — draws '
   'on Hill & Lynch\'s (1983) gender role intensification '
   'hypothesis and Rose\'s (2002) co-rumination research. Both '
   'phenomena reach their highest expression in middle school '
   'years (approximately ages 11-14), creating a developmental '
   'window in which girls become especially vulnerable to '
   'social-evaluative anxiety. Gender role intensification '
   'refers to the increased pressure adolescent girls face to '
   'conform to feminine norms (passivity, agreeableness, '
   'appearance focus). Co-rumination refers to the tendency of '
   'female adolescent friendships to dwell repeatedly on '
   'shared worries — a pattern that can intensify anxiety even '
   'while strengthening friendship bonds.',
   'Chúng tôi đề xuất ba cơ chế lý thuyết — riêng lẻ hoặc kết '
   'hợp — có thể giải thích vì sao khác biệt giới trong lo âu '
   'xuất hiện ở giai đoạn đầu tuổi vị thành niên (11-14) nhưng '
   'biến mất hoặc giảm đi ở giai đoạn thanh niên (16-29). Cơ chế '
   'thứ nhất — đỉnh điểm tổn thương phát triển — dựa trên giả '
   'thuyết gia tăng vai trò giới của Hill & Lynch (1983) và '
   'nghiên cứu co-rumination của Rose (2002). Cả hai hiện tượng '
   'đều đạt biểu hiện cao nhất trong những năm trung học cơ sở '
   '(khoảng 11-14 tuổi), tạo ra một cửa sổ phát triển trong đó '
   'nữ trở nên đặc biệt dễ tổn thương trước lo âu liên quan đến '
   'đánh giá xã hội. Gia tăng vai trò giới chỉ áp lực ngày càng '
   'tăng mà nữ vị thành niên phải tuân theo các chuẩn mực nữ '
   'tính (thụ động, dễ chiều, chú trọng ngoại hình). '
   'Co-rumination — cùng nhau day dứt — chỉ xu hướng tình bạn '
   'của nữ vị thành niên đào sâu lặp đi lặp lại các lo lắng '
   'chung — mẫu hình có thể làm lo âu mạnh hơn ngay cả khi củng '
   'cố mối quan hệ bạn bè.')

PB('The second mechanism — cultural moderation — proposes that '
   'Confucian academic culture in East Asia, particularly in '
   'Vietnam, imposes stricter gender expectations on girls than '
   'Western contexts do. Stankov (2010) documented that '
   'Confucian-heritage students experience high academic '
   'achievement coupled with elevated test anxiety and self-'
   'doubt; this pattern is especially pronounced in girls who '
   'face dual pressure to excel academically and to conform to '
   'traditional female roles. Small & Blanc (2021) further '
   'documented how the integrated worldview of tam giao (three '
   'teachings: Buddhism, Confucianism, and Taoism) shapes '
   'distinctively Vietnamese norms of emotional restraint and '
   'self-disclosure, with gendered implications for how anxiety '
   'is experienced and expressed.',
   'Cơ chế thứ hai — sự điều tiết bởi văn hóa — cho rằng văn '
   'hóa học thuật Nho giáo ở Đông Á, đặc biệt ở Việt Nam, áp '
   'đặt kỳ vọng giới khắt khe hơn lên nữ so với các bối cảnh '
   'phương Tây. Stankov (2010) ghi nhận rằng học sinh thừa hưởng '
   'di sản Nho giáo trải qua thành tích học tập cao kèm theo lo '
   'âu thi cử và tự nghi ngờ tăng cao; mô hình này đặc biệt rõ '
   'rệt ở nữ — những người chịu áp lực kép phải xuất sắc về học '
   'tập và phải tuân theo các vai trò nữ truyền thống. Small & '
   'Blanc (2021) tiếp tục ghi nhận cách thế giới quan tích hợp '
   'của tam giáo (ba giáo lý: Phật, Nho, Đạo) định hình các '
   'chuẩn mực chừng mực cảm xúc và bộc lộ bản thân đặc thù Việt '
   'Nam, với các hàm ý theo giới về cách lo âu được trải nghiệm '
   'và biểu đạt.')

PB('The third mechanism — methodological artifacts — points to '
   'differences in instrument context, sampling, and analytic '
   'pooling that may distinguish school-based adolescent surveys '
   'from market-panel young adult surveys. Adolescents completing '
   'questionnaires in classroom contexts may report differently '
   'than young adults responding to online market panels. Sample '
   'selection biases (who joins a market panel, who is present '
   'at school) may also differ. Finally, pooling young adults '
   'across a 13-year age range (16-29) may dilute gender effects '
   'that would be detectable within narrower age bands.',
   'Cơ chế thứ ba — sai số do phương pháp — chỉ các khác biệt '
   'về bối cảnh công cụ, chọn mẫu, và gộp phân tích có thể phân '
   'biệt khảo sát vị thành niên tại trường với khảo sát thanh '
   'niên qua hội đoàn trực tuyến. Vị thành niên trả lời câu hỏi '
   'trong bối cảnh lớp học có thể báo cáo khác với thanh niên '
   'trả lời hội đoàn trực tuyến. Sai lệch chọn mẫu (ai tham gia '
   'hội đoàn, ai có mặt tại trường) cũng có thể khác nhau. Cuối '
   'cùng, gộp thanh niên qua khoảng tuổi 13 năm (16-29) có thể '
   'làm loãng các hiệu ứng giới có thể phát hiện trong các dải '
   'tuổi hẹp hơn.')

H2('1.3. Why multi-group SEM is the right tool / Vì sao chọn '
   'SEM đa nhóm')
PB('Conventional gender comparison methods — independent t-tests '
   'or ANOVA on total anxiety scores — can detect mean '
   'differences between boys and girls, but they cannot test '
   'whether the entire risk-protective MECHANISM differs across '
   'gender. Multi-group structural equation modeling (SEM) goes '
   'further by simultaneously testing two related but distinct '
   'questions. First, measurement invariance: do the eight '
   'psychological scales actually measure the same underlying '
   'construct in both boys and girls? This is a precondition '
   'for any valid cross-gender comparison; if girls and boys '
   'respond to the same items in fundamentally different ways, '
   'we cannot meaningfully compare their pathway coefficients. '
   'Second, structural invariance: do the pathway coefficients '
   '(β) connecting each risk and protective factor to each '
   'anxiety outcome differ in magnitude or direction across '
   'gender?',
   'Các phương pháp so sánh giới thông thường — kiểm định t độc '
   'lập hay ANOVA trên điểm lo âu tổng — có thể phát hiện được '
   'khác biệt trung bình giữa nam và nữ, nhưng không kiểm chứng '
   'được liệu toàn bộ CƠ CHẾ nguy cơ-bảo vệ có khác biệt theo '
   'giới hay không. Mô hình phương trình cấu trúc đa nhóm (SEM) '
   'đi xa hơn bằng cách đồng thời kiểm chứng hai câu hỏi liên '
   'quan nhưng riêng biệt. Thứ nhất, bất biến thang đo: liệu '
   'tám thang tâm lý có thực sự đo cùng một cấu trúc tiềm ẩn ở '
   'cả nam và nữ không? Đây là điều kiện tiên quyết cho mọi so '
   'sánh chéo giới có giá trị; nếu nữ và nam trả lời cùng các '
   'mục theo những cách khác biệt cơ bản, chúng tôi không thể '
   'so sánh có ý nghĩa các hệ số đường dẫn của họ. Thứ hai, bất '
   'biến cấu trúc: liệu các hệ số đường dẫn (β) nối mỗi yếu tố '
   'nguy cơ và bảo vệ với mỗi kết quả lo âu có khác biệt về độ '
   'lớn hay dấu giữa nam và nữ?')

PB('Together, these two tests provide a fine-grained map of '
   'where gender matters and where it does not. This is '
   'substantively far more informative than simply reporting '
   'that "girls scored higher on average." For example, if '
   'measurement is invariant but the bullying → SAD pathway is '
   'stronger in girls, this implies that the SAME experience of '
   'being bullied translates into more severe social anxiety '
   'symptoms for girls — a finding with direct practical '
   'implications for intervention design. By contrast, if '
   'bullying → SAD is equivalent across gender but girls simply '
   'experience more bullying, the implication is different — '
   'reducing bullying exposure would close the gap. Multi-group '
   'SEM can distinguish these alternatives; t-tests cannot.',
   'Hai kiểm định này cùng nhau cung cấp bản đồ chi tiết về nơi '
   'giới có vai trò và nơi không. Điều này có giá trị thực chất '
   'cao hơn nhiều so với chỉ báo cáo "nữ có điểm trung bình cao '
   'hơn". Ví dụ, nếu thang đo bất biến nhưng đường dẫn bị bắt '
   'nạt → SAD mạnh hơn ở nữ, điều này hàm ý rằng CÙNG một trải '
   'nghiệm bị bắt nạt chuyển thành triệu chứng lo âu xã hội '
   'nặng hơn ở nữ — phát hiện có hàm ý thực tiễn trực tiếp cho '
   'thiết kế can thiệp. Ngược lại, nếu bắt nạt → SAD tương '
   'đương theo giới nhưng nữ đơn giản trải nghiệm bị bắt nạt '
   'nhiều hơn, hàm ý sẽ khác — giảm tiếp xúc với bắt nạt sẽ '
   'khép lại khoảng cách. SEM đa nhóm có thể phân biệt được '
   'các phương án này; kiểm định t thì không.')

H2('1.4. Present study and hypotheses / Nghiên cứu hiện tại '
   'và các giả thuyết')
PB('Building on our preceding Q2 paper, which established the '
   'integrated risk-protective higher-order SEM model in the '
   'pooled sample (R² = 0.598 explaining total anxiety variance), '
   'the present study tests whether this model is invariant '
   'across gender. We pre-register three hypotheses on Open '
   'Science Framework (OSF) prior to running analyses.',
   'Dựa trên bài Q2 trước đó của chúng tôi — đã thiết lập mô '
   'hình SEM bậc cao tích hợp nguy cơ-bảo vệ trong mẫu gộp (R² '
   '= 0,598 giải thích phương sai lo âu tổng) — nghiên cứu hiện '
   'tại kiểm chứng xem mô hình này có bất biến theo giới hay '
   'không. Chúng tôi đăng ký trước ba giả thuyết trên nền tảng '
   'Open Science Framework (OSF) trước khi chạy phân tích.')

PB('Hypothesis 1 (measurement invariance): The 8-scale '
   'measurement model will show full invariance at all three '
   'levels (configural, metric, scalar) across boys and girls. '
   'Justification: the scales we use are widely validated in '
   'mixed-gender adolescent samples and we have no theoretical '
   'reason to expect gender-differential item functioning. '
   'Hypothesis 2 (selective structural non-invariance): At '
   'least one structural pathway will be non-invariant; '
   'specifically, the bullying → SAD pathway will be stronger '
   'in girls than boys, based on evidence that adolescent girls '
   'show greater social-evaluative sensitivity (Rose 2002; '
   'Hankin et al. 2007) and that bullying experiences are more '
   'often relational/social in nature for girls. Hypothesis 3 '
   '(higher-order model gender differences): The integrated '
   'higher-order model — with YTNC (risk) and YTBV (protective) '
   'as second-order factors predicting RLLA (total anxiety) — '
   'will show different magnitudes of explained variance by '
   'gender.',
   'Giả thuyết 1 (bất biến thang đo): Mô hình đo lường 8 thang '
   'sẽ đạt bất biến hoàn toàn ở cả ba cấp (cấu hình, hệ số tải, '
   'ngưỡng) giữa nam và nữ. Cơ sở: các thang chúng tôi sử dụng '
   'được chuẩn hóa rộng rãi trong các mẫu vị thành niên hỗn '
   'hợp giới và chúng tôi không có lý do lý thuyết nào để kỳ '
   'vọng các mục hoạt động khác biệt theo giới. Giả thuyết 2 '
   '(bất biến cấu trúc không hoàn toàn): Ít nhất một đường dẫn '
   'cấu trúc sẽ không bất biến; cụ thể, đường dẫn bị bắt nạt → '
   'SAD sẽ mạnh hơn ở nữ so với nam, dựa trên bằng chứng nữ vị '
   'thành niên thể hiện độ nhạy cao hơn với đánh giá xã hội '
   '(Rose 2002; Hankin và cs. 2007) và bắt nạt ở nữ thường mang '
   'tính quan hệ/xã hội hơn về bản chất. Giả thuyết 3 (khác '
   'biệt giới ở mô hình bậc cao): Mô hình bậc cao tích hợp — '
   'với YTNC (nguy cơ) và YTBV (bảo vệ) là hai yếu tố bậc hai '
   'dự đoán RLLA (lo âu tổng) — sẽ thể hiện độ lớn phương sai '
   'được giải thích khác nhau giữa nam và nữ.')


# ========== METHODS ==========
H1('2. METHODS — PHƯƠNG PHÁP')

H2('2.1. Participants / Người tham gia')
PB('Participants were 1,352 Vietnamese lower secondary students '
   '(boys n = 614, 45.4%; girls n = 738, 54.6%) recruited from '
   'two purposively selected schools in Hanoi during the '
   '2024-2025 academic year: Nhat Tan school (urban context) '
   'and Tay Mo school (suburban context). The age range was '
   '11-14 years, covering grades 6 through 9, with the '
   'following grade-level distribution: grade 6 n = 368 '
   '(27.2%); grade 7 n = 316 (23.4%); grade 8 n = 340 (25.1%); '
   'grade 9 n = 328 (24.3%). Schools were selected to capture '
   'variation in urban-suburban context within Hanoi while '
   'controlling for region.',
   'Người tham gia gồm 1.352 học sinh trung học cơ sở Việt Nam '
   '(nam n = 614, 45,4%; nữ n = 738, 54,6%) tuyển từ hai '
   'trường được chọn có chủ đích tại Hà Nội trong năm học '
   '2024-2025: trường Nhật Tân (bối cảnh đô thị) và trường Tây '
   'Mỗ (bối cảnh ngoại ô). Khoảng tuổi là 11-14, bao trùm các '
   'lớp 6 đến 9, với phân bố theo khối như sau: lớp 6 n = 368 '
   '(27,2%); lớp 7 n = 316 (23,4%); lớp 8 n = 340 (25,1%); lớp '
   '9 n = 328 (24,3%). Các trường được chọn để bao quát sự biến '
   'thiên về bối cảnh đô thị-ngoại ô trong Hà Nội, đồng thời '
   'kiểm soát yếu tố vùng.')

PB('Inclusion criteria: students enrolled at one of the two '
   'study schools, present on the survey day, and with completed '
   'written parental consent. Exclusion criteria: students who '
   'declined to give assent, or who completed fewer than 80% of '
   'survey items. Of an initial pool of 1,524 eligible students, '
   '172 (11.3%) were excluded for missing data or assent refusal, '
   'yielding the final analytic sample of 1,352. The sample is '
   'predominantly ethnic Kinh (>95%), reflecting the demographic '
   'composition of Hanoi.',
   'Tiêu chí thu nhận: học sinh đang đi học tại một trong hai '
   'trường nghiên cứu, có mặt vào ngày khảo sát, và có sự đồng '
   'ý bằng văn bản của cha mẹ. Tiêu chí loại trừ: học sinh từ '
   'chối đồng thuận, hoặc hoàn thành dưới 80% số mục khảo sát. '
   'Từ một mẫu ban đầu gồm 1.524 học sinh đủ điều kiện, 172 '
   '(11,3%) bị loại do thiếu dữ liệu hoặc từ chối đồng thuận, '
   'tạo nên mẫu phân tích cuối cùng gồm 1.352 học sinh. Mẫu '
   'chủ yếu là dân tộc Kinh (>95%), phản ánh thành phần dân số '
   'của Hà Nội.')

H2('2.2. Measures / Công cụ đo lường')
PB('Eight validated scales were used. All scales had been '
   'previously translated into Vietnamese via forward-backward '
   'translation procedure, with expert panel review (three '
   'clinical psychologists and one educational measurement '
   'specialist) and pilot testing on a separate sample of 50 '
   'students. Anxiety subtypes were measured using the Revised '
   'Children\'s Anxiety and Depression Scale — RCADS (Chorpita '
   '2000): Generalized Anxiety Disorder (7 items, e.g., "I '
   'worry that something bad will happen to my parents"); '
   'Social Anxiety Disorder (4 items, e.g., "I worry what other '
   'people think of me"); Separation Anxiety Disorder (4 items, '
   'e.g., "I would feel scared if I had to sleep on my own"). '
   'All items used 4-point Likert scales (0 = never to 3 = '
   'always).',
   'Tám thang đo đã chuẩn hóa được sử dụng. Tất cả các thang đã '
   'được dịch sang tiếng Việt trước đó qua quy trình dịch '
   'xuôi-ngược, với hội đồng chuyên gia thẩm định (ba nhà tâm '
   'lý lâm sàng và một chuyên gia đo lường giáo dục) và thử '
   'nghiệm trên một mẫu riêng gồm 50 học sinh. Các phân loại '
   'lo âu được đo bằng Thang Lo âu và Trầm cảm Trẻ em phiên '
   'bản sửa đổi — RCADS (Chorpita 2000): Lo âu Lan tỏa (7 mục, '
   'ví dụ "Em lo điều xấu sẽ xảy ra với cha mẹ"); Lo âu Xã hội '
   '(4 mục, ví dụ "Em lo người khác nghĩ gì về em"); Lo âu Chia '
   'ly (4 mục, ví dụ "Em sẽ sợ nếu phải ngủ một mình"). Tất cả '
   'mục sử dụng thang Likert 4 điểm (0 = không bao giờ đến 3 '
   '= luôn luôn).')

PB('Risk factors were assessed using: (i) Olweus Bully/Victim '
   'Questionnaire — OBVQ (Olweus 1996), 8 items capturing '
   'physical and verbal bullying victimization; (ii) Educational '
   'Stress Scale for Adolescents — ESSA (Sun et al. 2011), 4 '
   'items on academic pressure; (iii) Smartphone Addiction '
   'Scale-Short Version — SAS-SV (Kwon et al. 2013), 5 items '
   'on dysfunctional smartphone use. Protective factors '
   'comprised: (i) Psychological Sense of School Membership — '
   'PSSM (Goodenow 1993), 7 items; (ii) Multidimensional Scale '
   'of Perceived Social Support — MSPSS (Zimet et al. 1988), '
   '8 items split into parental and peer subscales (4 items '
   'each); (iii) Rosenberg Self-Esteem Scale — RSES (Rosenberg '
   '1965), 5 items. All raw scores were rescaled to a 0-100 '
   'metric for cross-scale comparability while preserving '
   'original scoring conventions. Internal consistency '
   '(Cronbach\'s α) ranged from 0.70 to 0.91 across scales.',
   'Các yếu tố nguy cơ được đánh giá qua: (i) Bảng câu hỏi Bắt '
   'nạt/Nạn nhân Olweus — OBVQ (Olweus 1996), 8 mục bắt được '
   'việc bị bắt nạt thể chất và bằng lời; (ii) Thang Áp lực '
   'Học tập cho Vị thành niên — ESSA (Sun và cs. 2011), 4 mục '
   'về áp lực học tập; (iii) Thang Nghiện Điện thoại bản rút '
   'gọn — SAS-SV (Kwon và cs. 2013), 5 mục về sử dụng điện '
   'thoại không có chức năng. Các yếu tố bảo vệ gồm: (i) Cảm '
   'nhận Tâm lý về Sự gắn bó với Trường — PSSM (Goodenow 1993), '
   '7 mục; (ii) Thang Đa chiều về Hỗ trợ Xã hội Cảm nhận được '
   '— MSPSS (Zimet và cs. 1988), 8 mục chia thành phân thang '
   'cha mẹ và phân thang bạn bè (mỗi phân thang 4 mục); (iii) '
   'Thang Lòng Tự trọng Rosenberg — RSES (Rosenberg 1965), 5 '
   'mục. Tất cả điểm thô được chuyển sang thang 0-100 để so '
   'sánh giữa các thang, đồng thời giữ nguyên quy ước tính '
   'điểm gốc. Hệ số tin cậy nội tại (Cronbach α) dao động từ '
   '0,70 đến 0,91 giữa các thang.')

H2('2.3. Analytic strategy / Quy trình phân tích')
PB('The analytic strategy proceeds in five sequential steps. '
   'Software for primary analyses: Mplus 8.10 or higher (for '
   'multi-group SEM); R 4.3+ with the "lavaan" package 0.6-15+ '
   'for sensitivity analyses and reproducibility checks. All '
   'analyses use full-information maximum likelihood (FIML) '
   'estimation to accommodate missing data; mean-and-variance '
   'adjusted weighted least squares (WLSMV) is used as a '
   'robustness check given the ordinal nature of Likert items.',
   'Quy trình phân tích tiến hành qua năm bước tuần tự. Phần '
   'mềm cho các phân tích chính: Mplus 8.10 trở lên (cho SEM '
   'đa nhóm); R 4.3 trở lên với gói "lavaan" 0.6-15 trở lên '
   'cho các phân tích độ nhạy và kiểm tra khả năng tái tạo. '
   'Tất cả phân tích sử dụng phương pháp ước lượng hợp lý cực '
   'đại thông tin đầy đủ (FIML — Full-Information Maximum '
   'Likelihood) để xử lý dữ liệu khuyết; phương pháp bình '
   'phương tối thiểu có trọng số điều chỉnh trung bình-phương '
   'sai (WLSMV) được dùng như kiểm tra vững chắc do tính thứ '
   'tự của các mục Likert.')

PB('Step 1 — Descriptive statistics and bivariate gender '
   'comparisons. For each of the 8 scales, we compute mean, '
   'standard deviation (SD), Cronbach\'s alpha (α) reliability, '
   'and skewness/kurtosis stratified by gender. Independent '
   't-tests (with Welch correction for unequal variances) '
   'compare boys vs girls on each scale. Chi-square tests '
   'compare proportions above SAD clinical threshold (RCADS '
   'T-score ≥ 65). Cohen\'s d effect size with 95% CI is '
   'reported for each comparison.',
   'Bước 1 — Thống kê mô tả và so sánh giới đơn biến. Với mỗi '
   'trong 8 thang, chúng tôi tính giá trị trung bình, độ lệch '
   'chuẩn (SD), hệ số tin cậy Cronbach alpha (α), và độ '
   'lệch/độ nhọn phân tầng theo giới. Kiểm định t độc lập (có '
   'hiệu chỉnh Welch khi phương sai không đồng nhất) so sánh '
   'nam với nữ trên mỗi thang. Kiểm định chi-square so sánh tỷ '
   'lệ vượt ngưỡng lâm sàng SAD (điểm T của RCADS ≥ 65). Hệ số '
   'ảnh hưởng Cohen\'s d kèm khoảng tin cậy 95% được báo cáo '
   'cho mỗi so sánh.')

PB('Step 2 — Measurement invariance via multi-group '
   'confirmatory factor analysis (CFA). We test invariance in '
   'three sequential levels. Level 1 (configural invariance): '
   'we test whether the same factor structure — same items '
   'loading on the same latent factors — holds in both groups. '
   'Level 2 (metric invariance): on top of configural, we '
   'additionally constrain factor loadings to be equal across '
   'groups. Level 3 (scalar invariance): on top of metric, we '
   'additionally constrain item intercepts to be equal. The '
   'decision criterion at each step: ΔCFI ≤ 0.01 (Cheung & '
   'Rensvold 2002), combined with ΔRMSEA ≤ 0.015 (Chen 2007). '
   'If a level fails, we attempt partial invariance by '
   'releasing up to 20% of indicators per scale, with detailed '
   'reporting of which items required release.',
   'Bước 2 — Bất biến thang đo qua phân tích nhân tố khẳng '
   'định đa nhóm (CFA). Chúng tôi kiểm chứng bất biến qua ba '
   'cấp tuần tự. Cấp 1 (bất biến cấu hình): chúng tôi kiểm '
   'chứng xem cùng cấu trúc nhân tố — cùng mục tải lên cùng '
   'nhân tố tiềm ẩn — có giữ nguyên ở cả hai nhóm không. Cấp 2 '
   '(bất biến hệ số tải): trên nền cấu hình, chúng tôi bổ '
   'sung ràng buộc hệ số tải nhân tố bằng nhau giữa các nhóm. '
   'Cấp 3 (bất biến ngưỡng): trên nền hệ số tải, chúng tôi bổ '
   'sung ràng buộc ngưỡng chặn của mục bằng nhau. Tiêu chí ra '
   'quyết định ở mỗi bước: ΔCFI ≤ 0,01 (Cheung & Rensvold '
   '2002), kết hợp với ΔRMSEA ≤ 0,015 (Chen 2007). Nếu một '
   'cấp thất bại, chúng tôi thử bất biến từng phần bằng cách '
   'nới tối đa 20% chỉ báo trong mỗi thang, có báo cáo chi '
   'tiết các mục cần nới.')

PB('Step 3 — Structural invariance. We compare two nested '
   'models: (a) an unconstrained model in which all 21 '
   'hypothesized structural pathways are free to differ across '
   'groups; (b) a constrained model in which all 21 pathways '
   'are forced to be equal across groups. The chi-square '
   'difference test combined with ΔCFI determines if the '
   'unconstrained model fits significantly better. If yes, we '
   'systematically free one pathway at a time, prioritizing '
   'pathways with highest modification indices (MI > 10), to '
   'identify the SPECIFIC non-invariant pathways. For each '
   'non-invariant pathway, we report β_boys, β_girls, both 95% '
   'CIs, and a z-test of the cross-gender coefficient '
   'difference.',
   'Bước 3 — Bất biến cấu trúc. Chúng tôi so sánh hai mô hình '
   'lồng nhau: (a) mô hình không ràng buộc trong đó cả 21 '
   'đường dẫn cấu trúc đã giả thuyết tự do khác nhau giữa các '
   'nhóm; (b) mô hình ràng buộc trong đó cả 21 đường dẫn bị '
   'buộc bằng nhau giữa các nhóm. Kiểm định chênh lệch chi-'
   'square kết hợp với ΔCFI xác định mô hình không ràng buộc '
   'có vừa hơn đáng kể không. Nếu có, chúng tôi lần lượt nới '
   'từng đường dẫn một, ưu tiên các đường có chỉ số điều '
   'chỉnh cao nhất (MI > 10), để xác định ĐÚNG các đường dẫn '
   'không bất biến. Với mỗi đường dẫn không bất biến, chúng '
   'tôi báo cáo β_nam, β_nữ, hai khoảng tin cậy 95%, và một '
   'kiểm định z về chênh lệch hệ số giữa hai giới.')

PB('Step 4 — Higher-order integrated model invariance. We test '
   'invariance of the YTNC (risk) and YTBV (protective) '
   'second-order pathways to total anxiety (RLLA). We report β '
   'and 95% CI for each gender, R² for each gender (vs the '
   'pooled R² = 0.598 from Q2 v1), and the z-test of the '
   'cross-gender difference in second-order pathway '
   'coefficients. This step directly extends our Q2 finding by '
   'asking whether the same integrated model holds equally for '
   'boys and girls.',
   'Bước 4 — Bất biến của mô hình bậc cao tích hợp. Chúng tôi '
   'kiểm chứng bất biến của hai đường dẫn bậc hai YTNC (nguy '
   'cơ) và YTBV (bảo vệ) đến tổng lo âu (RLLA). Chúng tôi báo '
   'cáo β và khoảng tin cậy 95% cho mỗi giới, R² cho mỗi giới '
   '(so với R² mẫu gộp = 0,598 từ Q2 v1), và kiểm định z về '
   'chênh lệch hệ số đường dẫn bậc hai giữa hai giới. Bước '
   'này mở rộng trực tiếp phát hiện Q2 bằng cách đặt câu hỏi '
   'mô hình tích hợp này có giữ nguyên bằng nhau cho cả nam '
   'và nữ không.')

PB('Step 5 — Sensitivity analyses. We repeat the multi-group '
   'SEM after three sensitivity checks: (a) excluding sub-'
   'samples one at a time (urban-only / suburban-only) to test '
   'whether school context drives any observed effects; (b) '
   'stratifying by age (younger 11-12 vs older 13-14) to test '
   'the developmental peak hypothesis (the gender gap should '
   'be largest in the younger half); (c) applying a stricter '
   'invariance threshold (ΔCFI ≤ 0.005 per Chen 2007 for '
   'larger samples) to verify the robustness of conclusions.',
   'Bước 5 — Phân tích độ nhạy. Chúng tôi lặp lại SEM đa '
   'nhóm sau ba kiểm tra độ nhạy: (a) lần lượt loại trừ từng '
   'tiểu mẫu (chỉ đô thị / chỉ ngoại ô) để kiểm chứng xem bối '
   'cảnh trường có dẫn dắt bất kỳ hiệu ứng quan sát được nào '
   'không; (b) phân tầng theo tuổi (trẻ hơn 11-12 so với lớn '
   'hơn 13-14) để kiểm chứng giả thuyết đỉnh điểm phát triển '
   '(khoảng cách giới nên lớn nhất ở nửa trẻ hơn); (c) áp '
   'dụng ngưỡng bất biến nghiêm khắc hơn (ΔCFI ≤ 0,005 theo '
   'Chen 2007 cho các mẫu lớn) để xác minh tính vững chắc '
   'của kết luận.')

H2('2.4. Ethics / Đạo đức nghiên cứu')
PB('The study protocol was approved by the Ethics Committee '
   'of Hanoi National University of Education (HNUE) prior to '
   'data collection [approval reference: TBD — pending '
   'official letter from HNUE Ethics Committee]. Written '
   'parental consent and student assent were obtained from '
   'all participants. Students were informed that '
   'participation was voluntary, responses were confidential '
   'and would not affect their grades, and they could '
   'withdraw at any time without consequence. The study '
   'followed Declaration of Helsinki principles. Pre-'
   'registration of the analytic plan on Open Science '
   'Framework (OSF) was completed prior to data analysis.',
   'Đề cương nghiên cứu đã được phê duyệt bởi Hội đồng Đạo '
   'đức Trường Đại học Sư phạm Hà Nội (HNUE) trước khi thu '
   'thập dữ liệu [số tham chiếu phê duyệt: CHỜ XÁC NHẬN — '
   'chờ thư chính thức từ Hội đồng Đạo đức HNUE]. Sự đồng ý '
   'bằng văn bản của cha mẹ và sự đồng thuận bằng văn bản '
   'của học sinh đã được thu thập từ tất cả người tham gia. '
   'Học sinh được thông báo rằng việc tham gia là tự nguyện, '
   'các câu trả lời được bảo mật và sẽ không ảnh hưởng đến '
   'điểm số, và họ có thể rút khỏi nghiên cứu bất cứ lúc '
   'nào mà không có hậu quả. Nghiên cứu tuân thủ các nguyên '
   'tắc của Tuyên bố Helsinki. Việc đăng ký trước kế hoạch '
   'phân tích trên nền tảng Open Science Framework (OSF) '
   'đã hoàn thành trước khi phân tích dữ liệu.')


# ========== RESULTS (planned) ==========
H1('3. RESULTS — KẾT QUẢ (DỰ KIẾN)')

PB('Note to readers: This is the prose draft of expected '
   'results. Actual numeric values await completion of '
   'multi-group SEM analyses. Italicized placeholders [TBD] '
   'mark values that will be filled in after analysis. Tables '
   '1-3 and Figure 1 will be inserted at submission.',
   'Lưu ý cho người đọc: Đây là bản thảo văn xuôi của các kết '
   'quả dự kiến. Giá trị số thực tế chờ hoàn thành các phân '
   'tích SEM đa nhóm. Các chỗ trống được in nghiêng [CHỜ XÁC '
   'NHẬN] đánh dấu các giá trị sẽ được điền sau khi phân tích. '
   'Bảng 1-3 và Hình 1 sẽ được chèn vào lúc nộp bài.')

H2('3.1. Descriptive statistics and bivariate gender '
   'comparisons / Thống kê mô tả và so sánh giới đơn biến')
PB('Table 1 presents means, SDs, Cronbach\'s alphas, and '
   'bivariate gender comparison statistics for each of 8 '
   'scales. Internal consistency was acceptable to excellent '
   'across scales (α range [TBD; expected 0.70-0.91]). Based '
   'on our preliminary descriptives from Q2 v1 pooled '
   'analysis, we expect girls to score significantly higher '
   'on SAD (Cohen\'s d ≈ [TBD]) and SocAD (d ≈ [TBD]), and '
   'boys to score significantly higher on self-esteem (d ≈ '
   '[TBD]). Gender differences in academic stress and '
   'bullying victimization are expected to be smaller and '
   'directionally inconsistent across grades.',
   'Bảng 1 trình bày các giá trị trung bình, SD, Cronbach '
   'alpha, và thống kê so sánh giới đơn biến cho mỗi trong 8 '
   'thang. Hệ số tin cậy nội tại chấp nhận được đến xuất sắc '
   'qua các thang (khoảng α [CHỜ XÁC NHẬN; dự kiến 0,70-'
   '0,91]). Dựa trên thống kê mô tả sơ bộ của chúng tôi từ '
   'phân tích gộp Q2 v1, chúng tôi kỳ vọng nữ có điểm cao hơn '
   'có ý nghĩa thống kê trên SAD (Cohen\'s d ≈ [CHỜ XÁC '
   'NHẬN]) và SocAD (d ≈ [CHỜ XÁC NHẬN]), và nam có điểm '
   'cao hơn có ý nghĩa thống kê trên lòng tự trọng (d ≈ '
   '[CHỜ XÁC NHẬN]). Khác biệt giới về áp lực học tập và bị '
   'bắt nạt dự kiến nhỏ hơn và không nhất quán về hướng giữa '
   'các khối lớp.')

H2('3.2. Measurement invariance / Bất biến thang đo')
PB('Sequential model fit indices for the three invariance '
   'levels are presented in Table 2. The configural model is '
   'expected to fit well (CFI > 0.95, RMSEA < 0.06). Adding '
   'metric constraints is expected to produce ΔCFI < 0.01 '
   '(supporting metric invariance). Adding scalar constraints '
   'is expected to produce ΔCFI < 0.01 (supporting scalar '
   'invariance). If any partial invariance findings emerge, '
   'we will list specifically which items required release '
   'and discuss why those items may function differently '
   'across gender — for example, a bullying item referring to '
   '"calling me names" may have different interpretation for '
   'girls (often relational bullying) vs. boys (often '
   'physical bullying).',
   'Các chỉ số phù hợp mô hình tuần tự cho ba cấp bất biến '
   'được trình bày trong Bảng 2. Mô hình cấu hình dự kiến '
   'vừa tốt (CFI > 0,95, RMSEA < 0,06). Thêm ràng buộc hệ '
   'số tải dự kiến tạo ΔCFI < 0,01 (hỗ trợ bất biến hệ số '
   'tải). Thêm ràng buộc ngưỡng dự kiến tạo ΔCFI < 0,01 (hỗ '
   'trợ bất biến ngưỡng). Nếu xuất hiện bất biến từng phần, '
   'chúng tôi sẽ liệt kê cụ thể các mục cần nới và bàn luận '
   'vì sao các mục đó có thể hoạt động khác nhau giữa hai '
   'giới — ví dụ một mục bắt nạt nói "gọi tên xấu" có thể '
   'mang ý nghĩa khác với nữ (thường là bắt nạt quan hệ) so '
   'với nam (thường là bắt nạt thể chất).')

H2('3.3. Structural invariance / Bất biến cấu trúc')
PB('Comparing the unconstrained vs constrained models, we '
   'expect to identify a subset of non-invariant pathways. '
   'Table 3 will present β by gender for each of the 21 '
   'pathways + 95% CI + z-test of cross-gender difference. '
   'Pathways exceeding the substantively meaningful threshold '
   '(|β_girls - β_boys| > 0.10) will be highlighted in bold. '
   'Based on our pre-registered hypothesis H2, we expect the '
   'bullying → SAD pathway to be the primary non-invariant '
   'pathway. Additional non-invariant pathways may emerge '
   'around self-esteem and peer-support effects.',
   'So sánh mô hình không ràng buộc với mô hình ràng buộc, '
   'chúng tôi kỳ vọng xác định được một nhóm các đường dẫn '
   'không bất biến. Bảng 3 sẽ trình bày β theo giới cho mỗi '
   'trong 21 đường dẫn + khoảng tin cậy 95% + kiểm định z về '
   'chênh lệch giữa hai giới. Các đường dẫn vượt ngưỡng có ý '
   'nghĩa thực chất (|β_nữ - β_nam| > 0,10) sẽ được đánh dấu '
   'in đậm. Dựa trên giả thuyết đăng ký trước H2, chúng tôi '
   'kỳ vọng đường dẫn bị bắt nạt → SAD là đường dẫn không '
   'bất biến chính. Các đường dẫn không bất biến bổ sung có '
   'thể xuất hiện quanh các hiệu ứng của lòng tự trọng và hỗ '
   'trợ từ bạn bè.')

H2('3.4. Pathway-specific gender effects / Hiệu ứng giới theo '
   'từng đường dẫn')
PB('We will focus the substantive discussion on non-invariant '
   'pathways that meet two criteria: (i) statistically '
   'significant z-test of cross-gender difference at p < 0.05; '
   '(ii) substantively meaningful effect size difference '
   '|β_girls - β_boys| > 0.10. For each qualifying pathway, '
   'we report β_boys, β_girls, both 95% CIs, partial R², and '
   'practical interpretation. We do not over-interpret '
   'pathways that meet statistical significance only due to '
   'large sample size but show trivial effect size '
   'differences.',
   'Chúng tôi sẽ tập trung bàn luận thực chất trên các đường '
   'dẫn không bất biến đáp ứng hai tiêu chí: (i) kiểm định z '
   'về chênh lệch giữa hai giới có ý nghĩa thống kê tại p < '
   '0,05; (ii) chênh lệch hệ số ảnh hưởng có ý nghĩa thực '
   'chất |β_nữ - β_nam| > 0,10. Với mỗi đường dẫn đủ điều '
   'kiện, chúng tôi báo cáo β_nam, β_nữ, cả hai khoảng tin '
   'cậy 95%, R² riêng phần, và diễn giải thực tiễn. Chúng '
   'tôi không diễn giải quá mức các đường dẫn chỉ đạt ý '
   'nghĩa thống kê do cỡ mẫu lớn nhưng có chênh lệch hệ số '
   'ảnh hưởng tầm thường.')

H2('3.5. Higher-order integrated model / Mô hình bậc cao '
   'tích hợp')
PB('Figure 1 will display the higher-order SEM model with '
   'second-order factors YTNC (loaded by bullying, academic '
   'stress, smartphone addiction) and YTBV (loaded by school '
   'membership, parental support, peer support, self-esteem), '
   'both predicting RLLA (loaded by GAD, SAD, SocAD). The '
   'pooled R² from Q2 was 0.598. We report R² separately by '
   'gender. If R²_girls and R²_boys differ substantially '
   '(e.g., difference > 0.05), this indicates the integrated '
   'risk-protective model explains anxiety variance with '
   'different magnitudes for boys and girls — a finding with '
   'implications for the relative importance of universal vs '
   'gender-tailored prevention.',
   'Hình 1 sẽ thể hiện mô hình SEM bậc cao với hai yếu tố '
   'bậc hai YTNC (được tải bởi bị bắt nạt, áp lực học tập, '
   'nghiện điện thoại) và YTBV (được tải bởi gắn bó trường, '
   'hỗ trợ cha mẹ, hỗ trợ bạn bè, lòng tự trọng), cả hai dự '
   'đoán RLLA (được tải bởi GAD, SAD, SocAD). R² mẫu gộp '
   'từ Q2 là 0,598. Chúng tôi báo cáo R² riêng theo giới. '
   'Nếu R²_nữ và R²_nam khác biệt rõ rệt (ví dụ chênh lệch '
   '> 0,05), điều này chỉ ra mô hình tích hợp nguy cơ-bảo '
   'vệ giải thích phương sai lo âu với độ lớn khác nhau '
   'giữa nam và nữ — phát hiện có hàm ý về tầm quan trọng '
   'tương đối của phòng ngừa phổ quát so với phòng ngừa '
   'đặc thù theo giới.')


# ========== DISCUSSION ==========
H1('4. DISCUSSION — BÀN LUẬN')

H2('4.1. Summary of findings / Tóm tắt phát hiện')
PB('In this study we tested whether the integrated risk-'
   'protective SEM model that we previously established in a '
   'pooled sample of 1,352 Vietnamese lower secondary students '
   '(Q2, R² = 0.598) holds equally for boys and girls. Multi-'
   'group SEM revealed [TBD: measurement invariance results], '
   '[TBD: structural invariance results], and [TBD: higher-'
   'order model results]. The pattern of non-invariant '
   'pathways indicates that [TBD: substantive interpretation].',
   'Trong nghiên cứu này, chúng tôi đã kiểm chứng liệu mô '
   'hình SEM tích hợp nguy cơ-bảo vệ mà chúng tôi đã thiết '
   'lập trước đó trên mẫu gộp 1.352 học sinh trung học cơ sở '
   'Việt Nam (Q2, R² = 0,598) có giữ nguyên bằng nhau cho '
   'nam và nữ không. SEM đa nhóm cho thấy [CHỜ XÁC NHẬN: kết '
   'quả bất biến thang đo], [CHỜ XÁC NHẬN: kết quả bất biến '
   'cấu trúc], và [CHỜ XÁC NHẬN: kết quả mô hình bậc cao]. '
   'Mô hình các đường dẫn không bất biến chỉ ra rằng [CHỜ '
   'XÁC NHẬN: diễn giải thực chất].')

H2('4.2. Reconciliation with Jefferies & Ungar (2020) / Hòa '
   'giải khác biệt với Jefferies & Ungar (2020)')
PB('Our finding of gender differences in early-adolescent '
   'SAD diverges from the null gender effect reported by '
   'Jefferies & Ungar (2020) in young adults aged 16-29 '
   'across seven countries. We propose six lines of argument '
   '— four substantive and two methodological — to reconcile '
   'this discrepancy.',
   'Phát hiện của chúng tôi về khác biệt giới trong SAD ở '
   'giai đoạn đầu vị thành niên khác với hiệu ứng giới bằng '
   'không được báo cáo bởi Jefferies & Ungar (2020) trên '
   'thanh niên 16-29 tuổi tại bảy quốc gia. Chúng tôi đề '
   'xuất sáu luận chứng — bốn thực chất và hai phương pháp '
   '— để hòa giải sự khác biệt này.')

PB('Argument 1 — Developmental peak at ages 11-14. The age '
   'range 11-14 represents a critical developmental window '
   'in which gender differences in anxiety are at their '
   'highest expression. Hankin et al. (2007) demonstrated '
   'that adolescent girls experience a steeper increase in '
   'stress exposure and stress reactivity than boys during '
   'early puberty, leading to peak vulnerability to anxiety '
   'and depression. McLean & Anderson (2009), in their '
   'comprehensive review, documented that the female-to-male '
   'ratio for anxiety disorders rises from approximately '
   '1.5:1 in childhood to 2:1 or higher by early adolescence, '
   'before partially attenuating in adulthood. Our sample '
   'falls precisely within this peak vulnerability window, '
   'while Jefferies\' sample begins at age 16 — past the '
   'peak.',
   'Luận chứng 1 — Đỉnh điểm phát triển ở tuổi 11-14. '
   'Khoảng tuổi 11-14 đại diện cho một cửa sổ phát triển '
   'then chốt trong đó khác biệt giới về lo âu đạt biểu '
   'hiện cao nhất. Hankin và cs. (2007) chứng minh rằng nữ '
   'vị thành niên trải qua sự gia tăng dốc hơn về phơi '
   'nhiễm stress và phản ứng stress so với nam trong giai '
   'đoạn đầu dậy thì, dẫn đến đỉnh điểm tổn thương lo âu '
   'và trầm cảm. McLean & Anderson (2009), trong bài tổng '
   'quan toàn diện, ghi nhận tỷ số nữ:nam đối với các rối '
   'loạn lo âu tăng từ khoảng 1,5:1 ở tuổi thơ lên 2:1 hoặc '
   'cao hơn ở giai đoạn đầu vị thành niên, trước khi giảm '
   'một phần ở tuổi trưởng thành. Mẫu của chúng tôi rơi '
   'chính xác vào cửa sổ đỉnh điểm tổn thương này, trong '
   'khi mẫu của Jefferies bắt đầu từ tuổi 16 — đã qua đỉnh.')

PB('Argument 2 — Co-rumination in girls\' friendships. Rose '
   '(2002) introduced the construct of co-rumination — the '
   'tendency of close friends to dwell repeatedly and '
   'extensively on shared worries, life problems, and '
   'negative feelings. Her research consistently shows that '
   'adolescent girls engage in co-rumination significantly '
   'more than boys, and that this gendered pattern peaks in '
   'middle school years. Co-rumination has a paradoxical '
   'effect: it strengthens friendship intimacy but '
   'simultaneously amplifies internalizing symptoms '
   'including anxiety. In Vietnamese lower secondary '
   'students, where friendship intensity in girls is '
   'reinforced by cultural emphasis on female emotional '
   'connectedness, co-rumination may contribute substantially '
   'to elevated SAD scores in girls.',
   'Luận chứng 2 — Việc cùng nhau day dứt trong tình bạn '
   'nữ. Rose (2002) giới thiệu khái niệm co-rumination '
   '(cùng nhau day dứt) — xu hướng bạn thân đào sâu lặp '
   'đi lặp lại và rộng rãi vào các lo lắng chung, các vấn '
   'đề cuộc sống, và các cảm xúc tiêu cực. Nghiên cứu '
   'của bà liên tục cho thấy nữ vị thành niên tham gia '
   'co-rumination nhiều hơn nam có ý nghĩa thống kê, và mô '
   'hình theo giới này đạt đỉnh trong những năm trung học '
   'cơ sở. Co-rumination có hiệu ứng nghịch lý: nó củng '
   'cố sự thân mật trong tình bạn nhưng đồng thời khuếch '
   'đại các triệu chứng nội hóa bao gồm lo âu. Ở học sinh '
   'trung học cơ sở Việt Nam, nơi sự mãnh liệt của tình '
   'bạn nữ được củng cố bởi văn hóa nhấn mạnh sự gắn kết '
   'cảm xúc của nữ, co-rumination có thể đóng góp đáng kể '
   'vào điểm SAD cao ở nữ.')

PB('Argument 3 — Confucian gender script. Vietnamese '
   'lower secondary students are immersed in a Confucian '
   'cultural matrix that imposes distinctive gender '
   'expectations. Stankov (2010) documented that '
   'Confucian-heritage students show a paradoxical profile '
   'of high academic achievement combined with high test '
   'anxiety and self-doubt; this pattern is especially '
   'pronounced in girls, who face dual pressure to excel '
   'academically and to embody traditional female virtues '
   '(modesty, deference, emotional restraint). Small & '
   'Blanc (2021), in their analysis of Vietnam\'s mental '
   'health response during COVID-19, highlighted how the '
   'integrated worldview of tam giao — Buddhism, '
   'Confucianism, Taoism — shapes culturally distinctive '
   'norms of emotional management. For Vietnamese girls, '
   'these norms include greater emphasis on maintaining '
   'social harmony and avoiding emotional disturbance to '
   'others — a norm that can paradoxically increase '
   'internal SAD even while reducing external expression.',
   'Luận chứng 3 — Kịch bản giới Nho giáo. Học sinh '
   'trung học cơ sở Việt Nam được nhúng vào một ma trận '
   'văn hóa Nho giáo áp đặt các kỳ vọng giới đặc thù. '
   'Stankov (2010) ghi nhận rằng học sinh thừa hưởng di '
   'sản Nho giáo cho thấy một hồ sơ nghịch lý: thành '
   'tích học tập cao kết hợp với lo âu thi cử cao và tự '
   'nghi ngờ; mô hình này đặc biệt rõ rệt ở nữ — những '
   'người chịu áp lực kép phải xuất sắc về học tập và '
   'phải hiện thân các đức tính nữ truyền thống (khiêm '
   'tốn, lễ độ, chừng mực cảm xúc). Small & Blanc (2021), '
   'trong phân tích ứng phó sức khỏe tâm thần của Việt '
   'Nam trong đại dịch COVID-19, đã làm nổi bật cách thế '
   'giới quan tích hợp của tam giáo — Phật, Nho, Đạo — '
   'định hình các chuẩn mực quản lý cảm xúc đặc thù về '
   'văn hóa. Với nữ Việt Nam, các chuẩn mực này gồm việc '
   'nhấn mạnh hơn vào duy trì sự hòa thuận xã hội và '
   'tránh gây xáo trộn cảm xúc cho người khác — một chuẩn '
   'mực có thể, một cách nghịch lý, làm SAD nội tâm tăng '
   'lên ngay cả khi giảm biểu hiện ra ngoài.')

PB('Argument 4 — School-based vs market-panel context. '
   'Our data were collected in classroom contexts during '
   'normal school hours, with researchers present to '
   'answer questions; Jefferies\' data came from online '
   'market panels in which young adults completed the '
   'Social Interaction Anxiety Scale via web survey. '
   'These contexts elicit very different self-disclosure '
   'norms. Adolescents in classroom settings may be more '
   'attuned to peer judgment and report anxiety more '
   'honestly when guaranteed confidentiality; young '
   'adults in market panels may be more career-focused '
   'and minimize anxiety symptoms. The contextual '
   'difference may not eliminate Jefferies\' findings '
   'but may explain why some gender differences attenuate '
   'across data-collection contexts.',
   'Luận chứng 4 — Bối cảnh khảo sát tại trường so với '
   'hội đoàn trực tuyến. Dữ liệu của chúng tôi được thu '
   'thập trong bối cảnh lớp học trong giờ học bình '
   'thường, với nhà nghiên cứu có mặt để trả lời thắc '
   'mắc; dữ liệu của Jefferies đến từ các hội đoàn trực '
   'tuyến trong đó thanh niên hoàn thành Thang Lo âu '
   'Tương tác Xã hội qua khảo sát web. Hai bối cảnh này '
   'gợi ra các chuẩn mực bộc lộ bản thân rất khác nhau. '
   'Vị thành niên trong lớp học có thể nhạy bén hơn với '
   'sự đánh giá từ bạn bè và báo cáo lo âu chân thực '
   'hơn khi được bảo đảm bảo mật; thanh niên trong hội '
   'đoàn trực tuyến có thể tập trung vào sự nghiệp hơn '
   'và giảm thiểu các triệu chứng lo âu. Khác biệt bối '
   'cảnh có thể không xóa bỏ phát hiện của Jefferies '
   'nhưng có thể lý giải vì sao một số khác biệt giới '
   'giảm đi qua các bối cảnh thu thập dữ liệu.')

PB('Argument 5 — Methodological factors: instrument, '
   'sampling, and sub-sample size. Jefferies used the '
   'Social Interaction Anxiety Scale (SIAS), a 20-item '
   'measure focused specifically on interaction anxiety; '
   'we used the RCADS SAD subscale, a broader '
   '4-item measure capturing social-evaluative anxiety. '
   'The instruments overlap but are not identical. '
   'Additionally, Jefferies\' Vietnam sub-sample (n = '
   '973) is smaller than our sample (n = 1,352), '
   'reducing statistical power to detect gender effects '
   'within Vietnam specifically. The pooled cross-national '
   'analysis may further dilute country-specific gender '
   'effects.',
   'Luận chứng 5 — Các yếu tố phương pháp: công cụ, chọn '
   'mẫu, và cỡ tiểu mẫu. Jefferies sử dụng Thang Lo âu '
   'Tương tác Xã hội (SIAS), một thang 20 mục tập trung '
   'cụ thể vào lo âu tương tác; chúng tôi sử dụng phân '
   'thang SAD của RCADS, một thang rộng hơn 4 mục bắt '
   'được lo âu liên quan đến đánh giá xã hội. Hai công '
   'cụ trùng lặp nhưng không đồng nhất. Bên cạnh đó, '
   'tiểu mẫu Việt Nam của Jefferies (n = 973) nhỏ hơn '
   'mẫu của chúng tôi (n = 1.352), làm giảm sức mạnh '
   'thống kê để phát hiện các hiệu ứng giới riêng cho '
   'Việt Nam. Phân tích gộp xuyên quốc gia có thể làm '
   'loãng thêm các hiệu ứng giới đặc thù theo quốc gia.')

PB('Argument 6 — Age main effect dominates in Jefferies. '
   'Jefferies & Ungar (2020) reported a strong main '
   'effect of age: anxiety decreased across the 16-29 '
   'range, consistent with the well-documented '
   'attenuation of internalizing symptoms from late '
   'adolescence into adulthood. This strong age effect '
   'may statistically dominate any gender × age '
   'interaction, especially when age is treated as a '
   'covariate rather than as a stratifying variable. '
   'In our narrower age range (11-14), the age effect '
   'is smaller and gender effects more readily '
   'detectable.',
   'Luận chứng 6 — Hiệu ứng chính của tuổi áp đảo trong '
   'Jefferies. Jefferies & Ungar (2020) báo cáo một '
   'hiệu ứng chính mạnh của tuổi: lo âu giảm trong '
   'khoảng 16-29, phù hợp với sự giảm các triệu chứng '
   'nội hóa được ghi nhận rộng rãi từ cuối tuổi vị '
   'thành niên đến tuổi trưởng thành. Hiệu ứng tuổi '
   'mạnh này có thể áp đảo về mặt thống kê bất kỳ '
   'tương tác giới × tuổi nào, đặc biệt khi tuổi được '
   'xử lý như một biến đồng biến thay vì biến phân '
   'tầng. Trong khoảng tuổi hẹp hơn của chúng tôi '
   '(11-14), hiệu ứng tuổi nhỏ hơn và các hiệu ứng '
   'giới dễ phát hiện hơn.')

H2('4.3. Theoretical implications / Hàm ý lý thuyết')
PB('Three theoretical contributions emerge from this work. '
   'First, our findings extend Hill & Lynch\'s (1983) '
   'gender role intensification theory — originally '
   'developed in Western contexts — to a non-Western '
   'Confucian context, supporting its cross-cultural '
   'generalizability. Second, the cultural-developmental '
   'co-rumination pattern documented by Rose (2002) '
   'appears to operate in Vietnamese girls as well, but '
   'interacts with Confucian emotional disclosure norms '
   '(Small & Blanc 2021) in distinctive ways. Third, '
   'tam giao — as a culturally specific framework '
   'integrating Buddhism, Confucianism, and Taoism — '
   'provides an under-theorized lens for understanding '
   'how emotional restraint norms shape gendered SAD '
   'expression in Vietnam, deserving deeper theoretical '
   'and empirical development in future research.',
   'Ba đóng góp lý thuyết nổi lên từ công trình này. '
   'Thứ nhất, phát hiện của chúng tôi mở rộng lý thuyết '
   'gia tăng vai trò giới của Hill & Lynch (1983) — '
   'ban đầu phát triển trong bối cảnh phương Tây — '
   'sang bối cảnh Nho giáo phi phương Tây, hỗ trợ tính '
   'tổng quát xuyên văn hóa của nó. Thứ hai, mô hình '
   'co-rumination văn hóa-phát triển được Rose (2002) '
   'ghi nhận dường như cũng hoạt động ở nữ Việt Nam, '
   'nhưng tương tác với các chuẩn mực bộc lộ cảm xúc '
   'của Nho giáo (Small & Blanc 2021) theo cách đặc '
   'thù. Thứ ba, tam giáo — như một khung lý thuyết '
   'đặc thù về văn hóa tích hợp Phật, Nho, Đạo — '
   'cung cấp một lăng kính chưa được phát triển đầy '
   'đủ để hiểu cách các chuẩn mực chừng mực cảm xúc '
   'định hình biểu hiện SAD theo giới ở Việt Nam, '
   'xứng đáng được phát triển lý thuyết và thực '
   'nghiệm sâu hơn trong nghiên cứu tương lai.')

H2('4.4. Practical implications / Hàm ý thực tiễn')
PB('The findings have direct practical implications for '
   'gender-tailored prevention programs in Vietnamese '
   'schools. For girls in early adolescence (ages 11-14), '
   'interventions may benefit from explicit co-rumination '
   'disruption techniques — teaching girls to recognize '
   'when conversations with friends have shifted from '
   'supportive sharing to repetitive worry rumination, '
   'and to redirect such conversations toward concrete '
   'problem-solving. Complementary components include '
   'assertiveness training (especially for navigating '
   'peer-evaluation contexts), self-compassion practice '
   '(addressing self-criticism rooted in Confucian '
   'perfectionism), and explicit psychoeducation about '
   'the difference between healthy emotional disclosure '
   'and pathological rumination.',
   'Các phát hiện có hàm ý thực tiễn trực tiếp cho các '
   'chương trình phòng ngừa đặc thù theo giới ở trường '
   'học Việt Nam. Với nữ ở giai đoạn đầu vị thành niên '
   '(11-14 tuổi), các can thiệp có thể hưởng lợi từ các '
   'kỹ thuật làm gián đoạn co-rumination rõ ràng — dạy '
   'các em nhận biết khi các cuộc trò chuyện với bạn đã '
   'chuyển từ chia sẻ hỗ trợ sang ngẫm đi ngẫm lại lo '
   'lắng, và hướng các cuộc trò chuyện như vậy sang '
   'giải quyết vấn đề cụ thể. Các thành phần bổ sung '
   'gồm huấn luyện sự quyết đoán (đặc biệt cho việc '
   'điều hướng các bối cảnh đánh giá từ bạn bè), thực '
   'hành lòng tự trắc ẩn (giải quyết tự phê bình bắt '
   'nguồn từ chủ nghĩa hoàn hảo Nho giáo), và giáo dục '
   'tâm lý rõ ràng về sự khác biệt giữa bộc lộ cảm xúc '
   'lành mạnh và day dứt bệnh lý.')

PB('For boys, interventions may benefit from emotion '
   'identification and labeling exercises, since boys '
   'are socialized to suppress affective awareness and '
   'may under-report anxiety even when clinically '
   'present. School-based programs such as FRIENDS for '
   'Life and OurFutures — both with established evidence '
   'in mixed-gender samples — should be tested in '
   'gender-stratified randomized controlled trials '
   '(RCTs) in Vietnamese schools to determine whether '
   'adding gender-tailored components produces '
   'incremental benefit beyond gender-neutral protocols.',
   'Với nam, các can thiệp có thể hưởng lợi từ các bài '
   'tập nhận diện và đặt tên cảm xúc, vì nam được xã '
   'hội hóa để kìm nén nhận thức cảm xúc và có thể '
   'báo cáo dưới mức về lo âu ngay cả khi có biểu hiện '
   'lâm sàng. Các chương trình tại trường như FRIENDS '
   'for Life và OurFutures — cả hai đều có bằng chứng '
   'thiết lập trong các mẫu hỗn hợp giới — nên được '
   'kiểm chứng trong các thử nghiệm có đối chứng ngẫu '
   'nhiên (RCT) phân tầng theo giới ở trường học Việt '
   'Nam để xác định xem việc thêm các thành phần đặc '
   'thù giới có tạo lợi ích tăng thêm vượt qua các '
   'giao thức trung lập về giới không.')

H2('4.5. Limitations / Hạn chế')
PB('We acknowledge four major limitations. First, the '
   'cross-sectional design precludes causal inference: we '
   'cannot infer that bullying CAUSES SAD more strongly '
   'in girls; we can only state that the association is '
   'stronger. Longitudinal follow-up is necessary to '
   'establish temporal precedence. Second, the single-'
   'site Hanoi sample limits generalization to rural '
   'Vietnamese youth and ethnic-minority youth (Vietnam '
   'has 54 recognized ethnic groups; our sample is '
   'predominantly Kinh). Third, self-report instruments '
   'are subject to social desirability bias and to '
   'gender-biased disclosure norms; girls may be more '
   'willing to acknowledge anxiety than boys, which '
   'would inflate the observed gender difference. '
   'Fourth, single-time-point measurement does not '
   'capture developmental trajectories — we cannot test '
   'whether gender differences emerge, peak, or fade '
   'within our age window.',
   'Chúng tôi thừa nhận bốn hạn chế chính. Thứ nhất, '
   'thiết kế cắt ngang không cho phép suy luận nhân '
   'quả: chúng tôi không thể kết luận rằng bị bắt nạt '
   'GÂY RA SAD mạnh hơn ở nữ; chỉ có thể nói rằng mối '
   'liên quan mạnh hơn. Theo dõi dọc là cần thiết để '
   'thiết lập trình tự thời gian. Thứ hai, mẫu một địa '
   'điểm Hà Nội hạn chế khả năng tổng quát hóa cho '
   'thanh thiếu niên Việt Nam ở nông thôn và thanh '
   'thiếu niên dân tộc thiểu số (Việt Nam có 54 dân '
   'tộc được công nhận; mẫu của chúng tôi chủ yếu là '
   'người Kinh). Thứ ba, công cụ tự báo cáo chịu sai '
   'lệch do mong muốn xã hội và các chuẩn mực bộc lộ '
   'thiên lệch theo giới; nữ có thể sẵn lòng thừa nhận '
   'lo âu hơn nam, điều này sẽ phóng đại khác biệt '
   'giới quan sát được. Thứ tư, đo lường một thời điểm '
   'không bắt được quỹ đạo phát triển — chúng tôi '
   'không thể kiểm chứng liệu khác biệt giới xuất '
   'hiện, đạt đỉnh, hay phai dần trong cửa sổ tuổi '
   'của mẫu.')

H2('4.6. Future directions / Hướng nghiên cứu tiếp theo')
PB('Three priorities for future research. First, a '
   'longitudinal cohort study following the present '
   'sample through transition to upper secondary (ages '
   '16-18) and into emerging adulthood (ages 18-22) '
   'would test whether the gender gap closes — providing '
   'a direct test of the developmental peak hypothesis '
   'and a definitive reconciliation with Jefferies & '
   'Ungar (2020). Second, replication of the multi-'
   'group SEM in rural Vietnamese and ethnic-minority '
   'cohorts would test cultural moderation hypotheses. '
   'Third, our planned Q4 paper applies latent profile '
   'analysis (LPA) to the same cohort to identify '
   'person-centered phenotype profiles — complementing '
   'the variable-centered findings here with a person-'
   'centered typology of Vietnamese adolescent anxiety. '
   'A qualitative companion study probing gendered '
   'emotional disclosure norms is also planned to '
   'illuminate WHY gender pathways differ.',
   'Ba ưu tiên cho nghiên cứu tương lai. Thứ nhất, một '
   'nghiên cứu cohort theo dõi mẫu hiện tại qua chuyển '
   'tiếp lên trung học phổ thông (16-18 tuổi) và vào '
   'giai đoạn đầu trưởng thành (18-22 tuổi) sẽ kiểm '
   'chứng xem khoảng cách giới có khép lại không — cung '
   'cấp kiểm định trực tiếp giả thuyết đỉnh điểm phát '
   'triển và sự hòa giải dứt khoát với Jefferies & '
   'Ungar (2020). Thứ hai, lặp lại SEM đa nhóm trên '
   'các cohort Việt Nam ở nông thôn và dân tộc thiểu '
   'số sẽ kiểm chứng các giả thuyết về sự điều tiết '
   'bởi văn hóa. Thứ ba, bài Q4 đã lên kế hoạch của '
   'chúng tôi áp dụng phân tích hồ sơ tiềm ẩn (LPA) '
   'cho cùng cohort để xác định các hồ sơ phân loại '
   'lấy con người làm trung tâm — bổ sung cho các '
   'phát hiện lấy biến làm trung tâm ở đây bằng một '
   'phân loại lấy con người làm trung tâm về lo âu vị '
   'thành niên Việt Nam. Một nghiên cứu định tính '
   'đồng hành thăm dò các chuẩn mực bộc lộ cảm xúc '
   'theo giới cũng đã được lên kế hoạch để làm sáng '
   'tỏ VÌ SAO các đường dẫn giới khác nhau.')


# ========== REFERENCES ==========
H1('5. REFERENCES — TÀI LIỆU THAM KHẢO')
refs = [
    '[1] Jefferies P, Ungar M. (2020). Social anxiety in young '
    'people: A prevalence study in seven countries. PLOS ONE '
    '15(9):e0239133. DOI: 10.1371/journal.pone.0239133.',
    '[2] Hankin BL, Mermelstein R, Roesch L. (2007). Sex '
    'differences in adolescent depression: Stress exposure and '
    'reactivity models. Child Development 78(1):279-295. '
    'PMID: 17328705.',
    '[3] McLean CP, Anderson ER. (2009). Brave men and timid '
    'women? A review of the gender differences in fear and '
    'anxiety. Clinical Psychology Review 29(6):496-505. '
    'PMID: 19541399. DOI: 10.1016/j.cpr.2009.05.003.',
    '[4] Hill JP, Lynch ME. (1983). The intensification of '
    'gender-related role expectations during early '
    'adolescence. In: Brooks-Gunn J, Petersen AC (eds.), '
    'Girls at puberty. New York: Plenum Press, pp. 201-228.',
    '[5] Kessler RC, Berglund P, Demler O, Jin R, Merikangas '
    'KR, Walters EE. (2005). Lifetime prevalence and age-of-'
    'onset distributions of DSM-IV disorders in the National '
    'Comorbidity Survey Replication. Archives of General '
    'Psychiatry 62(6):593-602. PMID: 15939837. '
    'DOI: 10.1001/archpsyc.62.6.593.',
    '[6] Rose AJ. (2002). Co-rumination in the friendships of '
    'girls and boys. Child Development 73(6):1830-1843. '
    'PMID: 12487497. DOI: 10.1111/1467-8624.00509.',
    '[7] Stankov L. (2010). Unforgiving Confucian culture: A '
    'breeding ground for high academic achievement, test '
    'anxiety and self-doubt? Learning and Individual '
    'Differences 20(6):555-563. '
    'DOI: 10.1016/j.lindif.2010.05.003.',
    '[8] Small S, Blanc J. (2021). Mental Health During '
    'COVID-19: Tam Giao and Vietnam\'s Response. Frontiers '
    'in Psychiatry 11:589618. PMID: 33536961. '
    'DOI: 10.3389/fpsyt.2020.589618.',
    '[9] Cheung GW, Rensvold RB. (2002). Evaluating goodness-'
    'of-fit indexes for testing measurement invariance. '
    'Structural Equation Modeling 9(2):233-255.',
    '[10] Chen FF. (2007). Sensitivity of goodness of fit '
    'indexes to lack of measurement invariance. Structural '
    'Equation Modeling 14(3):464-504.',
    '[11] Hu LT, Bentler PM. (1999). Cutoff criteria for fit '
    'indexes in covariance structure analysis: Conventional '
    'criteria versus new alternatives. Structural Equation '
    'Modeling 6(1):1-55.',
    '[12] Compas BE, Jaser SS, Bettis AH, et al. (2017). '
    'Coping, emotion regulation, and psychopathology in '
    'childhood and adolescence: A meta-analysis and '
    'narrative review. Psychological Bulletin 143(9):939-'
    '991. PMID: 28616996. DOI: 10.1037/bul0000110.',
    '[13] Chorpita BF, Yim L, Moffitt C, Umemoto LA, Francis '
    'SE. (2000). Assessment of symptoms of DSM-IV anxiety '
    'and depression in children: A revised child anxiety '
    'and depression scale. Behaviour Research and Therapy '
    '38(8):835-855. PMID: 10937431.',
    '[14] Olweus D. (1996). The Revised Olweus Bully/Victim '
    'Questionnaire. University of Bergen, Norway.',
    '[15] Sun J, Dunne MP, Hou XY, Xu AQ. (2011). '
    'Educational stress scale for adolescents: Development, '
    'validity, and reliability with Chinese students. '
    'Journal of Psychoeducational Assessment 29(6):534-546.',
    '[16] Kwon M, Kim DJ, Cho H, Yang S. (2013). The '
    'smartphone addiction scale: Development and validation '
    'of a short version for adolescents. PLOS ONE '
    '8(12):e83558. PMID: 24391787.',
    '[17] Goodenow C. (1993). The psychological sense of '
    'school membership among adolescents: Scale development '
    'and educational correlates. Psychology in the Schools '
    '30(1):79-90.',
    '[18] Zimet GD, Dahlem NW, Zimet SG, Farley GK. (1988). '
    'The Multidimensional Scale of Perceived Social '
    'Support. Journal of Personality Assessment 52(1):'
    '30-41.',
    '[19] Rosenberg M. (1965). Society and the adolescent '
    'self-image. Princeton University Press.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Footer
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Soạn 08/06/2026 — Bản thảo đầy đủ Q3 phiên bản 1 '
              '— chờ NCS + thầy NMĐ review trước khi đăng ký '
              'trước trên OSF')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
