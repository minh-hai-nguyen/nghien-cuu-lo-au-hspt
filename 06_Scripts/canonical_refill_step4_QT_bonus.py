"""
STEP 4: Canonical 7 bài QT bonus (QT052-058) — BRIEF summary với cảnh báo scope.
Scope neurobiology/pharma lệch đề tài HS THCS/THPT, nhưng có giá trị tham chiếu cơ chế.
Japan Hikikomori (PDF corrupt) — skip với log.
"""
import sys, io, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
TT_DIR = BASE / 'Tom-tat-tung-bai'
FOLDER_AU = BASE / '02_Papers-goc' / 'The-gioi_Au-My-Uc'
FOLDER_KHAC = BASE / '02_Papers-goc' / 'The-gioi_Khac'

PAPERS = [
    {
        'id': 'QT052', 'folder_obj': FOLDER_AU,
        'descriptor': 'Mindfulness_NatureMH_IPD_MA_2023',
        'pdf_src': 'Mindfulness_Nature_MentalHealth_2023.pdf',
        'pdf_dst': 'QT052_Mindfulness_NatureMH_IPD_MA_2023.pdf',
        'title_vn': 'Tổng quan hệ thống và phân tích tổng hợp dữ liệu cá nhân các RCT về chương trình mindfulness',
        'title_en': 'Systematic review and individual participant data meta-analysis of randomized controlled trials assessing mindfulness-based programs',
        'journal': 'Nature Mental Health (2023) 1:462-476',
        'scope_warning': 'Đối tượng phần lớn là NGƯỜI LỚN. IPD MA giới hạn khi áp cho VTN — cần subgroup analysis.',
        'key_findings': 'IPD MA (individual participant data) — phương pháp cao nhất trong MA. RCT mindfulness cho MH. Tạp chí Nature Q1.',
        'scope_fit': 'MEDIUM — Mindfulness có thể adapt cho HS nhưng đa số RCT người lớn. Có tham chiếu.',
        'rating': 4,
    },
    {
        'id': 'QT053', 'folder_obj': FOLDER_AU,
        'descriptor': 'Pharmacotherapy_Anxiety_Review_Frontiers_2020',
        'pdf_src': 'Pharmacotherapy_Anxiety_2020_Frontiers.pdf',
        'pdf_dst': 'QT053_Pharmacotherapy_Anxiety_Review_Frontiers_2020.pdf',
        'title_vn': 'Tổng quan về điều trị dược lý rối loạn lo âu',
        'title_en': 'Pharmacotherapy for anxiety disorders (review)',
        'journal': 'Frontiers in Psychiatry (2020) 11:595584',
        'scope_warning': 'Đối tượng chủ yếu NGƯỜI LỚN. Với trẻ em/VTN, pharmacotherapy có guideline riêng (AACAP) — ít khuyến cáo ở độ tuổi nhỏ. CBT vẫn là first-line.',
        'key_findings': 'Review dược lý (SSRIs, SNRIs, benzodiazepines...). Là tài liệu ôn tập cơ chế + chỉ định lâm sàng.',
        'scope_fit': 'LOW — người lớn, không phải HS. Chỉ tham khảo khi cần hiểu mechanism dược lý.',
        'rating': 3,
    },
    {
        'id': 'QT054', 'folder_obj': FOLDER_AU,
        'descriptor': 'ChronicStress_Neuroinflammation_FrontPsych_2023',
        'pdf_src': 'Chronic_Stress_Neuroinflammation_FrontPsych_2023.pdf',
        'pdf_dst': 'QT054_ChronicStress_Neuroinflammation_FrontPsych_2023.pdf',
        'title_vn': 'Stress mạn tính, viêm thần kinh và trầm cảm: Tổng quan cơ chế bệnh sinh',
        'title_en': 'Chronic stress, neuroinflammation, and depression: an overview of pathophysiological mechanisms and emerging anti-inflammatories',
        'journal': 'Frontiers in Psychiatry (2023)',
        'scope_warning': 'Review cơ chế SINH HỌC CƠ BẢN (neuroinflammation pathway) — không về can thiệp/dịch tễ HS. Đối tượng chung (không riêng VTN).',
        'key_findings': 'Viêm thần kinh là cơ chế chung giữa stress mạn và trầm cảm. Emerging anti-inflammatory targets.',
        'scope_fit': 'LOW — basic science, không áp dụng trực tiếp đề tài lo âu HS.',
        'rating': 3,
    },
    {
        'id': 'QT055', 'folder_obj': FOLDER_AU,
        'descriptor': 'Parade_Epigenetics_ChildhoodMaltreatment_TranslPsych_2021',
        'pdf_src': 'Epigenetics_Childhood_Maltreatment_TranslPsych_2021.pdf',
        'pdf_dst': 'QT055_Parade_Epigenetics_Maltreatment_TranslPsych_2021.pdf',
        'title_vn': 'Tổng quan hệ thống về ngược đãi thời thơ ấu và methyl hóa DNA: gene ứng viên và nghiên cứu toàn bộ bộ gen',
        'title_en': 'A systematic review of childhood maltreatment and DNA methylation: candidate gene and genome-wide studies',
        'authors': 'Parade et al.',
        'journal': 'Translational Psychiatry (2021) 11:134',
        'scope_warning': 'Review EPIGENETICS — cơ chế phân tử. Có liên quan ACEs (đã note trong VN015 Ngô Anh Vinh 2024) nhưng không về can thiệp HS.',
        'key_findings': 'DNA methylation là cơ chế epigenetic liên kết ACEs với rối loạn MH về sau. Gene ứng viên: NR3C1, SLC6A4...',
        'scope_fit': 'MEDIUM — có thể dùng làm nền cho "vì sao ACEs ảnh hưởng lâu dài" trong thảo luận VN015.',
        'rating': 3,
    },
    {
        'id': 'QT056', 'folder_obj': FOLDER_KHAC,
        'descriptor': 'Jiang_Microbiota_GutBrain_Anxiety_FrontNeuro_2024',
        'pdf_src': 'Microbiota_GutBrain_Anxiety_2024_FrontNeuro.pdf',
        'pdf_dst': 'QT056_Jiang_Microbiota_GutBrain_Anxiety_FrontNeuro_2024.pdf',
        'title_vn': 'Giao tiếp qua trục microbiota-gut-brain trong lo âu',
        'title_en': 'Microbiota-gut-brain axis communication in anxiety',
        'authors': 'Jiang M, Kang L, Wang YL, Zhou B, Li HY, Yan Q, Liu ZG',
        'journal': 'Frontiers in Neuroscience (2024)',
        'scope_warning': 'Review cơ chế gut-brain axis — hướng mới nhưng chưa có can thiệp lâm sàng trường học. Đối tượng chung.',
        'key_findings': 'Hệ vi sinh đường ruột ảnh hưởng não bộ qua trục microbiota-gut-brain; có thể liên quan lo âu.',
        'scope_fit': 'LOW — basic science, không áp dụng đề tài.',
        'rating': 3,
    },
    {
        'id': 'QT057', 'folder_obj': FOLDER_KHAC,
        'descriptor': 'NeuralCircuits_Anxiety_FrontNeuralCirc_2025',
        'pdf_src': 'Neural_Circuits_Mechanisms_Anxiety_2025_FrontNeuralCirc.pdf',
        'pdf_dst': 'QT057_NeuralCircuits_Mechanisms_Anxiety_FrontNeuralCirc_2025.pdf',
        'title_vn': 'Tiến độ nghiên cứu về cơ chế mạch thần kinh của lo âu',
        'title_en': 'Research progress on the neural circuits mechanisms of anxiety',
        'journal': 'Frontiers in Neural Circuits (2025) · Beijing University of Chinese Medicine',
        'scope_warning': 'Review về MẠCH THẦN KINH (amygdala, BNST, PFC) — basic neuroscience, không về dịch tễ/can thiệp HS.',
        'key_findings': 'Amygdala + bed nucleus stria terminalis (BNST) + prefrontal cortex là các mạch chính của lo âu. Review tiến độ 2025.',
        'scope_fit': 'LOW — basic neuroscience.',
        'rating': 3,
    },
    {
        'id': 'QT058', 'folder_obj': FOLDER_KHAC,
        'descriptor': 'Guo_Neuroinflammation_Neuromodulation_TranslPsych_2022',
        'pdf_src': 'Neuroinflammation_Neuromodulation_TranslPsych_2022.pdf',
        'pdf_dst': 'QT058_Guo_Neuroinflammation_Neuromodulation_TranslPsych_2022.pdf',
        'title_vn': 'Cơ chế viêm thần kinh của các liệu pháp kích thích điều biến thần kinh (neuromodulation) cho lo âu và trầm cảm',
        'title_en': 'Neuroinflammation mechanisms of neuromodulation therapies for anxiety and depression',
        'authors': 'Guo B, Zhang M, Hao W, Wang Y, Zhang T, Liu C',
        'journal': 'Translational Psychiatry (2022)',
        'scope_warning': 'Review cơ chế — neuromodulation (TMS, tDCS, DBS) là can thiệp xâm lấn, KHÔNG phù hợp can thiệp trường học.',
        'key_findings': 'Neuromodulation (TMS, tDCS) giảm neuroinflammation. Cơ chế cho can thiệp y học kháng trị.',
        'scope_fit': 'LOW — can thiệp xâm lấn, không phù hợp can thiệp trường học VN.',
        'rating': 3,
    },
]

# Handle Japan Hikikomori (PDF corrupt)
HIKIKOMORI_SRC = FOLDER_KHAC / 'Japan_Youth_Suicide_Hikikomori_2025.pdf'

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def build_brief_docx(paper, out_path):
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    t = d.add_heading(f"{paper['id']} — Tóm tắt ngắn (bài bonus)", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)

    # Scope warning box
    warn = d.add_table(rows=1, cols=1)
    warn.style = 'Table Grid'
    wc = warn.rows[0].cells[0]
    shade_cell(wc, 'FCE4D6')
    p = wc.paragraphs[0]
    r = p.add_run('⚠ CẢNH BÁO SCOPE: ')
    r.bold = True
    r.font.color.rgb = RGBColor(192, 0, 0)
    p.add_run(paper['scope_warning'])
    d.add_paragraph()

    # Metadata
    meta_tbl = d.add_table(rows=5, cols=2)
    meta_tbl.style = 'Table Grid'
    meta_rows = [
        ('Tên (VN)', paper['title_vn']),
        ('Tên (EN)', paper['title_en']),
        ('Tác giả', paper.get('authors', '(xem PDF)')),
        ('Nơi công bố', paper['journal']),
        ('Mức độ phù hợp đề tài', paper['scope_fit']),
    ]
    for i, (k, v) in enumerate(meta_rows):
        cell0 = meta_tbl.rows[i].cells[0]
        shade_cell(cell0, 'D9E1F2')
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        meta_tbl.rows[i].cells[1].text = v
    d.add_paragraph()

    # Key findings
    h = d.add_heading('Phát hiện/kết luận chính', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)
    d.add_paragraph(paper['key_findings'])
    d.add_paragraph()

    # Note for future
    h = d.add_heading('Khi nào dùng', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(191, 97, 14)
    d.add_paragraph(
        f"Bài này thuộc nhóm BONUS (basic science/review ngoài scope HS THCS/THPT). "
        f"Chỉ dùng khi cần tham chiếu cơ chế sinh học/dược lý/neuromodulation. "
        f"KHÔNG dùng trực tiếp cho khuyến nghị can thiệp trường học VN. "
        f"Nếu cần tóm tắt đầy đủ Jenkins template, cần đọc kỹ PDF và mở rộng."
    )

    stars = '⭐' * paper['rating']
    h = d.add_heading(f"Đánh giá: {stars} ({paper['rating']}/5 — chất lượng NC; riêng scope-fit đã ghi trên)", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(90, 90, 90)

    d.save(out_path)

log = []
with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

for paper in PAPERS:
    src = paper['folder_obj'] / paper['pdf_src']
    dst = paper['folder_obj'] / paper['pdf_dst']
    if not src.exists():
        log.append(f"SKIP {paper['id']}: source {paper['pdf_src']} missing")
        continue
    if not dst.exists():
        src.rename(dst)
        log.append(f"RENAMED {paper['id']}: {dst.name}")
    tt_name = f"{paper['id']}_{paper['descriptor']}.docx"
    tt_path = TT_DIR / tt_name
    build_brief_docx(paper, tt_path)
    log.append(f"SUMMARY {paper['id']}: {tt_name} ({tt_path.stat().st_size//1024} KB)")

    idx[paper['id']] = {
        'id': paper['id'],
        'descriptor': paper['descriptor'],
        'summary': tt_name,
        'translation': None,
        'pdf': paper['pdf_dst'],
        'pdf_folder': paper['folder_obj'].name,
        'scope': 'bonus_out_of_scope',
    }
    log.append(f"INDEX {paper['id']}: added (scope=bonus)")

# Handle Japan Hikikomori (corrupt PDF)
if HIKIKOMORI_SRC.exists():
    log.append(f"\nQT059 Japan Hikikomori 2025: PDF corrupt (Stream ended unexpectedly)")
    log.append(f"  → Skipping canonical. PDF kept as-is at {HIKIKOMORI_SRC.name}")
    log.append(f"  → Cần tải lại PDF từ nguồn gốc trước khi canonical.")

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

n_vn = sum(1 for k in idx if k.startswith('VN'))
n_qt = sum(1 for k in idx if k.startswith('QT'))
log.append(f"\nTotal canonical entries: {len(idx)} ({n_vn} VN + {n_qt} QT)")
for l in log:
    print(l)
