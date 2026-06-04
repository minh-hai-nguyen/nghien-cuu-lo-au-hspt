# -*- coding: utf-8 -*-
"""Khoi phuc tu BEFORE_TOC va them lai TOC SACH (khong placeholder).
28/05/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC_BACKUP = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026_BEFORE_TOC.docx')
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

# 1. Restore from BEFORE_TOC backup (clean state)
shutil.copy(SRC_BACKUP, FILE)
print(f"Restored from: {SRC_BACKUP}")

# 2. Open and add TOC field WITHOUT placeholder text
d = Document(FILE)

# Find MỤC LỤC
toc_idx = None
for i, p in enumerate(d.paragraphs[:200]):
    if p.text.strip().upper() in ('MỤC LỤC', 'MUC LUC'):
        toc_idx = i
        break

if toc_idx is None:
    print("KHONG TIM THAY MỤC LỤC!")
    sys.exit(1)
print(f"Found MỤC LỤC at para {toc_idx}")


def make_toc_paragraph_clean(parent_para):
    """TOC field WITHOUT placeholder text. Word will fill it on F9."""
    r = parent_para.add_run()
    # Begin
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fldChar_begin)
    # Instruction
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r._r.append(instrText)
    # Separate
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fldChar_sep)
    # NO placeholder text - just end immediately
    # End
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar_end)


# Insert new paragraph after MỤC LỤC heading
toc_heading_p = d.paragraphs[toc_idx]
new_p_xml = OxmlElement('w:p')
toc_heading_p._element.addnext(new_p_xml)
toc_p_obj = Paragraph(new_p_xml, toc_heading_p._parent)
make_toc_paragraph_clean(toc_p_obj)

# Save
d.save(FILE)
print(f"Saved: {FILE}")

# Verify
d2 = Document(FILE)
has_toc = False
for elem in d2.element.body.iter():
    if elem.tag == qn('w:instrText'):
        if elem.text and 'TOC' in elem.text:
            has_toc = True
            print(f"✓ TOC field intact: {elem.text!r}")
            break
if not has_toc:
    print("✗ TOC MISSING!")

# Check new paragraph (61) text - should be empty
if 61 < len(d2.paragraphs):
    p61 = d2.paragraphs[61]
    print(f"Para 61 text: {p61.text!r}")
    if p61.text.strip() == '':
        print("✓ Khong co placeholder text — sach")
    else:
        print(f"⚠ Co text: {p61.text}")

print(f"\nSize: {os.path.getsize(FILE)//1024} KB")
print(f"Paragraphs: {len(d2.paragraphs)}")
