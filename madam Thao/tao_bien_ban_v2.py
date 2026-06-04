# -*- coding: utf-8 -*-
"""
Tạo 2 file DOCX biên bản cuộc họp từ file .md đã có dấu tiếng Việt đầy đủ.
Căn độ rộng cột bảng theo nội dung.
"""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_col_width(cell, width_cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(width_cm * 567)))
    w.set(qn('w:type'), 'dxa')
    tcW.append(w)

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def make_doc(include_stories):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)

    def h(text, level=1):
        hd = doc.add_heading(text, level=level)
        for r in hd.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0,0,0)

    def p(text, bold=False, italic=False):
        para = doc.add_paragraph()
        r = para.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = bold
        r.italic = italic

    # Read content from md file (already has proper Vietnamese diacritics)
    with open('Bien ban hop - Madam Thao.md', 'r', encoding='utf-8') as f:
        md = f.read()

    # TITLE
    suffix = ' (Bản đầy đủ)' if include_stories else ' (Bản tóm tắt)'
    title = doc.add_heading('', level=0)
    r = title.add_run('BIÊN BẢN CUỘC HỌP' + suffix + '\nChương trình đào tạo lãnh đạo — Hội nhập & Tỏa sáng')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p('Chủ trì: TS. Nguyễn Thị Phương Thảo — Chủ tịch HĐQT Vietjet', bold=True)
    p('Thành phần: Anh Tuấn và các thành viên nhóm đào tạo, truyền thông')
    p('Nội dung: Xây dựng chương trình đào tạo lãnh đạo kết hợp hội nhập tỏa sáng cho Vietjet')
    doc.add_paragraph()

    # I
    h('I. Tinh thần cốt lõi của chương trình đào tạo')
    p('Madam Thảo nhấn mạnh rằng chương trình đào tạo phải đi vào chiều sâu, không chỉ dựa trên kịch bản có sẵn. Mục tiêu là tạo môi trường để người tham gia tìm thấy giá trị bản thân trong tập thể. Tinh thần ban đầu của các chương trình đào tạo tại Vietjet đã bị mai một — thiếu yếu tố con người, thiếu "hồn".')
    p('Chị muốn trở lại chương trình đào tạo với điểm khác biệt nằm ở yếu tố con người: khi gặp khó khăn, tập trung vào cách vượt qua và phối hợp đội nhóm, chứ không chỉ làm cho xong.')

    # II
    h('II. Thực trạng nhân sự')
    p('Madam Thảo chia sẻ thẳng thắn: tinh thần nhân viên Vietjet những năm qua đã giảm sút đáng kể. Người ta rời đi không phải vì chế độ (chế độ tốt), mà vì môi trường — không tìm thấy ý nghĩa công việc.')
    p('Hình ảnh ngày xưa: tổ bay briefing tại trụ sở công ty, xếp hàng ngay ngắn, diễu trong văn phòng — nhân viên nhìn thấy tổ bay chuẩn bị làm nhiệm vụ thì tự hào, thấy giá trị và ý nghĩa. Bây giờ: đầu tắt mặt tối, ra khỏi nhà 2-3 giờ sáng, ra sân bay rồi về đi ngủ, không có thời gian tiêu, không thấy ý nghĩa gắn bó.')

    # III
    h('III. Kế hoạch chương trình cụ thể')
    p('Quy mô: 60 người, chia 5 team', bold=True)
    p('Bối cảnh: Chương trình lồng ghép với khóa học 6 tháng "Tỏa sáng Tương lai". Làm 2 trong 1: 1 tiếng đầu dành cho tinh thần Vietjet, phần còn lại là hội nhập tỏa sáng dành cho lãnh đạo.')

    p('Ngày 1:', bold=True)
    p('• Buổi sáng: Khởi động với truyền thống Vietjet, truyền lửa (~1 tiếng); tiếp theo hội nhập tỏa sáng cho lãnh đạo')
    p('• Nội dung lãnh đạo: công tác lập kế hoạch, quản lý tài chính, quản lý nhân sự, ngân sách, chỉ tiêu, chỉ số')
    p('• Buổi chiều: Nội dung chuyên sâu + giới thiệu chính thức')
    p('• Tối: Gala Dinner, team building (tập luyện từ tối hôm trước)')

    p('Ngày 2:', bold=True)
    p('• Đưa vào các nội dung kỹ năng chuyên sâu cho lãnh đạo')
    p('• Tham khảo và chắt lọc từ chương trình của 3 đối tác bên ngoài: A11 (tư vấn lãnh đạo), INSEAD (giáo dục điều hành), và một đối tác thứ ba')
    p('• Chủ tịch nhấn mạnh: chương trình đào tạo CEO phức tạp, không đơn giản ngồi tự xây — cần tài liệu chuẩn quốc tế, tham khảo rồi chắt lọc thành chương trình riêng phù hợp Vietjet')

    p('Team building:', bold=True)
    p('• Chia thành 5 đội, mỗi đội có đội trưởng phụ trách. Bắt đầu tập luyện từ tối thứ 6 (trước ngày khai mạc)')
    p('• Mỗi đội chuẩn bị 01 tiết mục trình diễn tại Gala Dinner buổi tối')
    p('• Tổ chức trò chơi tập thể (team game) trong hội trường — không gian đủ rộng để triển khai')
    p('• Thiết kế đồng phục riêng cho mỗi đội phục vụ thi đấu')
    p('• Kinh nghiệm từ các sự kiện tại Nha Trang và Galaxy cho thấy tinh thần tham gia rất tích cực, bao gồm cả lãnh đạo cấp cao (ông Đặng đã tham gia diễn kịch tại Galaxy)')

    p('Truyền thông:', bold=True)
    p('• Đoàn truyền thông đi cùng từ đầu, ghi nhận mọi khoảnh khắc')

    # IV - Table with proper widths
    h('IV. Phân công công việc')
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [4.0, 12.0]  # Người: 4cm, Nhiệm vụ: 12cm
    headers = ['Người', 'Nhiệm vụ']
    rows_data = [
        ('Anh Tuấn', 'Chuẩn bị Agenda chương trình theo hướng dẫn của Chủ tịch HĐQT. Nắm giữ và tổng hợp tài liệu từ các đối tác.'),
        ('Nhóm đào tạo', 'Ghi nhận nội dung cuộc họp. Gửi lại toàn bộ tài liệu và nội dung đã thực hiện cho nhóm.'),
        ('Nhóm truyền thông', 'Đồng hành ghi nhận hình ảnh xuyên suốt sự kiện. Xây dựng clip về Chủ tịch HĐQT phục vụ chương trình.'),
        ('Chủ tịch HĐQT', 'Chốt kịch bản chương trình 2 ngày. Chắt lọc nội dung từ tài liệu đối tác để xây dựng chương trình riêng Vietjet.'),
    ]
    for row in t.rows:
        for ci in range(2):
            set_col_width(row.cells[ci], widths[ci])
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for para in c.paragraphs:
            for r in para.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shade(c, 'D9E2F3')
    for ri, (k, v) in enumerate(rows_data):
        t.rows[ri+1].cells[0].text = k
        t.rows[ri+1].cells[1].text = v
        for ci in range(2):
            for para in t.rows[ri+1].cells[ci].paragraphs:
                for r in para.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    doc.add_paragraph()

    # V
    h('V. Định hướng lãnh đạo — Chủ tịch HĐQT')
    p('Nguyên tắc nhân quả trong quản lý: Chủ tịch nhấn mạnh triết lý nhân quả — kết quả công việc phản ánh chất lượng đầu vào. Xử lý vấn đề nội bộ cần thận trọng, tránh sử dụng biện pháp hành chính căng thẳng. Mỗi cá nhân có quyền được tạo cơ hội khắc phục khi mắc sai sót.')
    p('Xây dựng sự công bằng trong tổ chức: Sự công bằng đòi hỏi thời gian xây dựng cơ chế đánh giá rõ ràng, không dựa trên quyết định cảm tính của cá nhân. Chủ tịch lưu ý rằng nếu lãnh đạo cấp trên không gương mẫu trong khi áp dụng kỷ luật đối với cấp dưới thì không đảm bảo tính công bằng.')
    p('Phát triển năng lực coaching: Mỗi lãnh đạo cần trở thành người hướng dẫn (coach) cho đội ngũ của mình. Tinh thần học tập và đào tạo phải được duy trì liên tục, trân trọng và không để mai một.')
    p('Bổ sung kỹ năng quản lý: Chủ tịch chỉ ra rằng tổ chức đang thiếu một số kỹ năng quản lý cơ bản như lập kế hoạch, quản lý dự án. Đây không phải vấn đề khó mà là chưa được đào tạo bài bản. Khóa đào tạo này cần cập nhật những kiến thức này một cách có hệ thống.')

    # VI - Stories (full version only)
    if include_stories:
        h('VI. Câu chuyện truyền cảm hứng — Chủ tịch HĐQT chia sẻ')

        h('1. Nâng công suất sân bay Tân Sơn Nhất', 2)
        p('"Năm 2013 mình làm, 2014 bắt đầu chạy. Dự án nâng cao năng lực điều hành cất hạ cánh ở sân bay Tân Sơn Nhất — nâng lên gấp đôi từ 30 lượt đến 60 lượt. Bây giờ đang vận hành 48 lượt. Từ 24 triệu lượt khách lên 42 triệu lượt khách. Dự án ấy làm cho cả ngành hàng không. ACV mười mấy năm liền, mỗi năm đón 16 triệu lượt khách, lợi nhuận mấy trăm triệu đô la một năm. Chứ không phải là doanh số đâu. Thì làm sao mà chẳng có tiền đi xây sân bay."')

        h('2. Giải cứu sàn chứng khoán', 2)
        p('"Giải cứu cho cả cái thị trường vốn. Cả nước Mỹ có sàn New York, Nasdaq — 2 cái sàn chứng khoán. Cả nước Anh có London Stock Exchange. Mà nó bị treo thì giải cứu mấy trăm tỷ đô la ở đấy. Có hơn 100 công ty chứng khoán, 4 triệu khách hàng là tổ chức và cá nhân trong nước và ngoài nước. Mỗi ngày giao dịch khoảng 1 tỷ đô la, mà cao là 2 tỷ đô la. Mình cũng nhảy vào giải cứu."')
        p('"Bao nhiêu người được hưởng lợi từ đó. Chính phủ thì được cái sàn, được cả nước. Các công ty, các nhà đầu tư thì được hưởng lợi. Bản thân các công ty niêm yết ở đấy thì người ta có thị trường vốn để huy động vốn, hoạt động, phát triển. Thế còn mình thì không cần cái gì cả."')

        h('3. Ngoại giao qua mua máy bay', 2)
        p('"Mua máy bay Boeing là đi với nước Mỹ — ông Trump, ông Obama, ông Bush. Mua máy bay Airbus là đi với Pháp, Đức, Tây Ban Nha — những nước sản xuất Airbus. Mua máy bay COMAC của Trung Quốc là chương trình của ông Tập Cận Bình. Là mình tạo nên xây dựng mối quan hệ ngoại giao, xây dựng kinh tế giữa các quốc gia, dân tộc. Mình sẽ giảm bớt đi những chiến tranh, những sự xung đột."')

        h('4. APEC và năng lực triển khai thần tốc', 2)
        p('"APEC là 2 lần mình xây dựng, năm 2017. Đưa Việt Nam vào thế giới và đưa thế giới về Việt Nam. Cái giải pháp chiếu sáng mình làm có 90 ngày thôi — trong khi đó Hàn Quốc triển khai 12 năm mà không chạy, đến lúc bàn giao đã 17 năm. Xây cung hội nghị 4 tháng."')
        p('"Ngày 9 tháng 1 năm ngoái, không ai mời mình đâu, nhưng mà thấy thời cuộc nó căng thẳng. Mình có mối quan hệ tốt với ông Trump, Đại sứ quán Mỹ mời mình, thì mình bay luôn, mang máy bay đi luôn. Cũng tạo nên thiện cảm, cũng đỡ căng thẳng."')

        h('5. Thiện nguyện COVID và triết lý thiết kế', 2)
        p('"Mùa Covid mình lao đi chở vắc-xin, thiết bị y tế, bao nhiêu những bao gạo. Không có những cái đó thì làm sao mà người ta được yên lòng."')
        p('"Tất cả những cái gì mà mình làm — thiết kế một cái nhà chung cư, là phải có cái thang máy lớn để dành cho cáng cứu thương và quan tài. Từ cái thiết kế là mình cũng nghĩ đến con người — lúc người ta bệnh, người ta cần phải đưa đi cáng. Ở đâu có bậc thang là mình phải có cái đường dốc cho xe nôi, cho xe người khuyết tật. Đấy là tinh thần mình đưa vào mọi hoạt động."')

        h('6. Đóng góp toàn diện', 2)
        p('"Cờ vua lên top 7 thế giới. Bóng đá nam vào World Cup, bóng đá nữ cũng vào World Cup. Toàn diện cho đất nước, ở tất cả các lĩnh vực — văn hóa, xã hội, kinh tế, tài chính."')
        p('Chủ tịch HĐQT nhấn mạnh: "Mỗi con người đều có thể trở nên vĩ đại. Làm những việc lớn có ý nghĩa, mà từ những hoạt động rất là đơn giản, giản dị của mình thôi."', italic=True)

    # VII / VI - Kết luận
    step = 'VII' if include_stories else 'VI'
    h(f'{step}. Kết luận và bước tiếp theo')
    p('1. Chốt kịch bản chương trình 2 ngày')
    p('2. Chuẩn bị tài liệu đào tạo — tham khảo từ đối tác, chắt lọc với chị')
    p('3. Lồng ghép chương trình hội nhập tỏa sáng với sự kiện Vietjet')
    p('4. Phân chia 5 team, chuẩn bị team building từ tối hôm trước')
    p('5. Gửi lại tài liệu và nội dung đã làm cho nhóm')
    p('6. Xây dựng clip về Madam Thảo để phục vụ chương trình')
    p('7. Viết lại câu chuyện Madam Thảo — cần người biết kể chuyện, có tư liệu')

    return doc

# Generate both versions
doc1 = make_doc(include_stories=True)
doc1.save('Biên bản họp - Đầy đủ - Madam Thảo.docx')
print('Done: Biên bản họp - Đầy đủ - Madam Thảo.docx')

doc2 = make_doc(include_stories=False)
doc2.save('Biên bản họp - Tóm tắt - Madam Thảo.docx')
print('Done: Biên bản họp - Tóm tắt - Madam Thảo.docx')
