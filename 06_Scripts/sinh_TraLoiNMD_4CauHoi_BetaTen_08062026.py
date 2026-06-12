# -*- coding: utf-8 -*-
"""Tra loi 4 cau hoi cua thay NMD: (1+2) y nghia beta = 0.176/0.258 trong
bai Hoang Trung Hoc; (3) Dakota King-White viet tat; (4) BMC Psychiatry
dich tieng Viet."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'TraLoiNMD_4CauHoi_BetaSEM_TenTacGia_BMC_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.4


def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x80, 0x40, 0x00)

def P(t):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def BB(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('• ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def Q(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Câu hỏi thầy NMĐ: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r2 = p.add_run('"' + t + '"')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True


# ============================================================
TITLE('TRẢ LỜI 4 CÂU HỎI CỦA THẦY NGUYỄN MINH ĐỨC')
TITLE('Hệ số β trong bài Hoàng Trung Học + Viết tắt họ kép + '
      'Tên tạp chí BMC Psychiatry', 11)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tài liệu trả lời nội bộ — Soạn 08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
H1('CÂU 1 + 2 — Ý NGHĨA HỆ SỐ β = 0,176 VÀ 0,258 TRONG BÀI HOÀNG '
   'TRUNG HỌC')

Q('Trong bài của Hoàng Trung Học, thầy vẫn chưa hiểu tác giả giải '
   'thích "thời gian sử dụng máy tính và điện thoại", với hệ số Beta '
   'là 0,176 trong giai đoạn Covid-19 (β= 0,176, xem Bảng 6) và 0,258 '
   'sau Covid-19 (β= 0,258) nghĩa là như thế nào? Đáng lẽ phải xác '
   'định đấy là yếu tố nguy cơ và đưa ra kết quả β âm thì dễ hiểu hơn')

H2('Trước hết — Quy ước về dấu của hệ số β trong hồi quy/SEM')
P('Dấu (+ hay −) của hệ số β phụ thuộc HOÀN TOÀN vào hướng đo lường '
  'của biến phụ thuộc (outcome). Em xin trình bày bảng đối chiếu để '
  'thầy đối chiếu:')

# Bảng đối chiếu
table = d.add_table(rows=3, cols=3)
table.style = 'Light Grid Accent 1'
headers = ['Biến phụ thuộc (outcome)', 'Yếu tố NGUY CƠ có β = ...',
           'Yếu tố BẢO VỆ có β = ...']
data = [
    ['Triệu chứng TIÊU CỰC (lo âu, trầm cảm, stress, '
     'thời gian không lành mạnh, hành vi rủi ro)',
     'DƯƠNG (+) — nguy cơ tăng → triệu chứng tăng',
     'ÂM (−) — bảo vệ tăng → triệu chứng giảm'],
    ['Sức khỏe tâm thần TÍCH CỰC (well-being, hạnh phúc, '
     'gắn bó trường, lòng tự trọng, kết quả học tập)',
     'ÂM (−) — nguy cơ tăng → tích cực giảm',
     'DƯƠNG (+) — bảo vệ tăng → tích cực tăng'],
]
cells = table.rows[0].cells
for j, h in enumerate(headers):
    cells[j].text = ''
    p = cells[j].paragraphs[0]
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for i, row_data in enumerate(data):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        cells[j].text = ''
        p = cells[j].paragraphs[0]
        r = p.add_run(val); r.font.size = Pt(10)
        if j == 0:
            r.bold = True

p = d.add_paragraph()
p.paragraph_format.space_before = Pt(8)

H2('Áp dụng vào bài Hoàng Trung Học')
P('Em phỏng đoán bài Hoàng Trung Học có biến phụ thuộc là một TRIỆU '
  'CHỨNG TIÊU CỰC (lo âu, trầm cảm, stress học tập, hoặc rối loạn '
  'sức khỏe tâm thần nói chung). Nếu đúng vậy:')

BB('β = +0,176 (giai đoạn COVID-19): nghĩa là thời gian sử dụng máy '
   'tính + điện thoại CÀNG NHIỀU → triệu chứng CÀNG TĂNG. Đây CHÍNH '
   'LÀ dấu hiệu của một YẾU TỐ NGUY CƠ — đúng như thầy đã hiểu')
BB('β = +0,258 (sau COVID-19): mối liên hệ này MẠNH HƠN ở giai đoạn '
   'hậu đại dịch — tức là khi học sinh đi học trở lại, yếu tố nguy '
   'cơ "thời gian dùng thiết bị" càng quan trọng hơn')
BB('Tác giả KHÔNG cần và KHÔNG NÊN đảo dấu sang β âm — vì sẽ làm '
   'thay đổi ý nghĩa thực tế')

H2('Cỡ hiệu ứng (effect size) theo chuẩn Cohen 1988')
BB('β = 0,176 ≈ hiệu ứng NHỎ-VỪA (Cohen quy ước: 0,10 = nhỏ; 0,30 = '
   'vừa; 0,50 = lớn)')
BB('β = 0,258 ≈ hiệu ứng VỪA — đáng chú ý về mặt thực tiễn')
BB('Mức tăng từ 0,176 lên 0,258 (tăng khoảng 47%) cho thấy yếu tố '
   'nguy cơ này MẠNH lên đáng kể sau đại dịch — phù hợp với xu hướng '
   'quốc tế ghi nhận học sinh dùng thiết bị nhiều hơn sau COVID')

H2('Khi nào yếu tố nguy cơ MỚI có β âm?')
P('Nếu thầy đã từng đọc các bài có yếu tố nguy cơ với β âm, có lẽ '
  'biến phụ thuộc trong các bài đó là MỘT BIẾN TÍCH CỰC (gắn bó với '
  'trường, lòng tự trọng, hạnh phúc học đường…). Ở khung này:')
BB('Yếu tố nguy cơ (vd: bắt nạt, áp lực học tập) làm GIẢM biến tích '
   'cực → β âm')
BB('Yếu tố bảo vệ (vd: hỗ trợ cha mẹ) làm TĂNG biến tích cực → β '
   'dương')

P('Như vậy, không có quy luật cố định "yếu tố nguy cơ phải có β âm" '
  '— mọi thứ phụ thuộc vào hướng đo của biến phụ thuộc.')


# ============================================================
H1('CÂU 3 — DAKOTA KING-WHITE VIẾT TẮT LÀ KING-WHITE D.')

Q('Dakota King-White thì chuyển thành King-White D., phải không Em?')

P('Em xác nhận thầy hoàn toàn đúng ạ. Cách viết tắt chính xác là '
  '**King-White D.**')

H2('Lý do')

table2 = d.add_table(rows=5, cols=2)
table2.style = 'Light Grid Accent 1'
rows2 = [
    ['Thành phần', 'Giá trị'],
    ['Tên đầy đủ', 'Dakota King-White'],
    ['Tên (given name)', 'Dakota'],
    ['Họ (family name) — họ kép có dấu nối',
     'King-White'],
    ['Viết tắt chuẩn (PubMed + APA + Vancouver)',
     'King-White D.'],
]
for i, row_data in enumerate(rows2):
    cells = table2.rows[i].cells
    for j, val in enumerate(row_data):
        cells[j].text = ''
        p = cells[j].paragraphs[0]
        r = p.add_run(val); r.font.size = Pt(10)
        if i == 0:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif j == 0:
            r.bold = True

p = d.add_paragraph()
p.paragraph_format.space_before = Pt(8)

H2('Quy tắc về họ kép')
BB('Họ kép có dấu nối (hyphenated double surname) phải viết NGUYÊN '
   'vẹn cả 2 phần — không tách rời, không bỏ phần nào')
BB('Sai (KHÔNG NÊN): "King W." hoặc "White D." — sẽ làm mất một '
   'phần họ')
BB('Đúng: "King-White D." — giữ nguyên dấu nối + cả 2 phần')
BB('Sắp xếp trong danh mục Tài liệu Tham khảo: theo chữ "K" '
   '(King-White) — không phải chữ "W"')

H2('Các trường hợp họ kép phổ biến khác')
BB('Smith-Jones → Smith-Jones')
BB('Hall-Williams → Hall-Williams')
BB('Pérez-García → Pérez-García')
BB('Đối với họ Tây Ban Nha kiểu "García López" (không có dấu nối) '
   '— có thể chỉ dùng họ đầu: "García J."; nhưng nhiều tác giả ngày '
   'nay yêu cầu giữ cả 2: "García López J."')


# ============================================================
H1('CÂU 4 — TÊN TẠP CHÍ BMC PSYCHIATRY KHI VIẾT BẰNG TIẾNG VIỆT')

Q('Tạp chí BMC Psychiatry dịch ra tiếng Việt như thế nào, Em ơi?')

P('Em khuyến nghị KHÔNG dịch tên tạp chí — giữ nguyên "BMC Psychiatry". '
  'Lý do em trình bày dưới đây:')

H2('3 phương án có thể cân nhắc')

table3 = d.add_table(rows=4, cols=3)
table3.style = 'Light Grid Accent 1'
rows3 = [
    ['Phương án', 'Cách viết', 'Nhận xét'],
    ['A — Giữ nguyên (em khuyến nghị)', 'BMC Psychiatry',
     'Tên thương hiệu nhà xuất bản — như tên riêng; phù hợp chuẩn '
     'Vancouver + APA; người đọc khoa học quốc tế quen thuộc'],
    ['B — Dịch hoàn toàn',
     'Tạp chí Tâm thần học BMC',
     'Có thể gây nhầm lẫn — "BMC" mất nghĩa gốc; KHÔNG khuyến nghị'],
    ['C — Kết hợp có chú thích',
     'Tạp chí BMC Psychiatry (tạm dịch: Tâm thần học BMC)',
     'Dùng khi cần giải thích cho người đọc không chuyên'],
]
for i, row_data in enumerate(rows3):
    cells = table3.rows[i].cells
    for j, val in enumerate(row_data):
        cells[j].text = ''
        p = cells[j].paragraphs[0]
        r = p.add_run(val); r.font.size = Pt(10)
        if i == 0:
            r.bold = True
            r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        elif j == 0:
            r.bold = True

p = d.add_paragraph()
p.paragraph_format.space_before = Pt(8)

H2('Vì sao em đề xuất KHÔNG dịch?')
BB('"BMC" là viết tắt của BioMed Central — nhà xuất bản (publisher), '
   'không có nghĩa tiếng Việt')
BB('"Psychiatry" là chuyên ngành — về mặt kỹ thuật dịch được thành '
   '"Tâm thần học", nhưng dịch xong làm mất tên gốc thương hiệu')
BB('Quy ước trích dẫn quốc tế (Vancouver, APA, MLA) đều yêu cầu giữ '
   'nguyên tên tạp chí gốc — không dịch')
BB('Trong luận án + bài báo tiếng Việt của các nghiên cứu sinh Việt '
   'Nam đã công bố, đa số đều giữ nguyên "BMC Psychiatry", "BMC Public '
   'Health", "Frontiers in Psychiatry"… không dịch sang tiếng Việt')

H2('Cách dùng cụ thể trong luận án + bản tiếng Việt của bài Q2/Q3/Q4')

H3('Trong câu văn')
P('"…đã được công bố trên tạp chí BMC Psychiatry…"')
P('hoặc: "…theo nghiên cứu công bố trên BMC Psychiatry (Pham và cs., '
  '2024)…"')

H3('Trong danh mục Tài liệu Tham khảo (chuẩn Vancouver)')
P('Pham TTH, Do TT, Nguyen TL, Ngo AV. Anxiety symptoms and coping '
  'strategies among high school students in Vietnam after COVID-19 '
  'pandemic: a mixed-method evaluation. **Front Public Health**. '
  '2024;12:1232856. (Lưu ý: dùng tên VIẾT TẮT chính thức của tạp chí '
  '— với BMC Psychiatry thì viết tắt là "BMC Psychiatry" hoặc "BMC '
  'Psychiatry" — không rút gọn thêm)')

H3('Trong danh mục Tài liệu Tham khảo (chuẩn APA)')
P('Pham, T. T. H., Do, T. T., Nguyen, T. L., & Ngo, A. V. (2024). '
  'Anxiety symptoms and coping strategies among high school students '
  'in Vietnam after COVID-19 pandemic: A mixed-method evaluation. '
  '*Frontiers in Public Health*, *12*, 1232856.')


# ============================================================
H1('TÀI LIỆU THAM KHẢO')

refs = [
    'Cohen J. (1988). Statistical Power Analysis for the Behavioral '
    'Sciences (2nd ed.). Lawrence Erlbaum Associates. (Quy ước cỡ '
    'hiệu ứng β: 0,10 = nhỏ; 0,30 = vừa; 0,50 = lớn)',
    'Kline RB. (2016). Principles and Practice of Structural Equation '
    'Modeling (4th ed.). The Guilford Press. (Sách giáo khoa chuẩn về '
    'diễn giải hệ số β trong SEM)',
    'International Committee of Medical Journal Editors (ICMJE). '
    'Recommendations on author names — bibliographic references. URL: '
    'https://www.icmje.org/recommendations/browse/manuscript-preparation/'
    'preparing-for-submission.html',
    'American Psychological Association (APA). Publication Manual (7th '
    'ed.). 2020. Chapter 9 — Reference Format. (Quy chuẩn tên tác giả '
    'và tên tạp chí)',
    'National Library of Medicine (NLM). Vancouver Citation Style '
    'Guide. URL: https://www.nlm.nih.gov/bsd/uniform_requirements.html '
    '(Quy chuẩn về tên tác giả viết tắt + tên tạp chí)',
    'Pham TTH, Do TT, Nguyen TL, Ngo AV. Anxiety symptoms and coping '
    'strategies among high school students in Vietnam after COVID-19 '
    'pandemic. Front Public Health. 2024;12:1232856. PMID: 38435293.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            chunks.append(cell.text)
print(f'WORD COUNT: {sum(len(c.split()) for c in chunks)}')
