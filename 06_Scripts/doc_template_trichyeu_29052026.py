# -*- coding: utf-8 -*-
"""Doc template Trich yeu luan an cua Tran Anh Khoi de hieu cau truc.
29/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LATS = os.path.join(ROOT, 'Luận án TS')

TEMPLATE = None
for f in os.listdir(LATS):
    # Tran Anh Khoi file specifically - starts with "Tr" Vietnamese
    if f.lower().endswith('.docx') and 'Anh' in f and 'Kh' in f and ('ch' in f or 'tr' in f.lower()):
        TEMPLATE = os.path.join(LATS, f)
        print(f'Found: {f}')
        break

d = Document(TEMPLATE)

print()
print('=== ALL PARAGRAPHS ===')
for i, para in enumerate(d.paragraphs, 1):
    if para.text.strip():
        font_info = ''
        if para.runs:
            r = para.runs[0]
            sz = r.font.size.pt if r.font.size else '?'
            font_info = f' [sz={sz} b={r.bold} i={r.italic} fn={r.font.name}]'
        print(f'[{i}]{font_info} a={para.alignment} | {para.text}')

print()
print('=== ALL TABLES ===')
for ti, tbl in enumerate(d.tables, 1):
    print(f'--- Table {ti} ({len(tbl.rows)}r x {len(tbl.columns)}c) ---')
    for ri, row in enumerate(tbl.rows):
        for ci, cell in enumerate(row.cells):
            txt = cell.text.strip()
            if txt:
                # Limit long text
                preview = txt[:200] + '...' if len(txt) > 200 else txt
                print(f'  R{ri}C{ci}: {preview}')

print()
print('=== SECTIONS / MARGINS ===')
for si, sec in enumerate(d.sections, 1):
    print(f'S{si}: W={sec.page_width.cm:.2f} H={sec.page_height.cm:.2f}')
    print(f'  T={sec.top_margin.cm:.2f} B={sec.bottom_margin.cm:.2f} '
          f'L={sec.left_margin.cm:.2f} R={sec.right_margin.cm:.2f}')

print()
print('=== STYLE Normal ===')
s = d.styles['Normal']
sz = s.font.size.pt if s.font.size else '?'
print(f'Font: {s.font.name}, Size: {sz}')
print(f'Line spacing: {s.paragraph_format.line_spacing}')
