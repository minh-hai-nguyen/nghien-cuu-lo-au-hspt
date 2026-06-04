# -*- coding: utf-8 -*-
"""Add 4 remaining images (p001 logos) to VN002 FULL."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_imgs/'

doc = Document(OUT)

# Find the first page marker paragraph
anchor = None
for p in doc.paragraphs:
    if p.text.strip().startswith('--- Trang 7'):
        anchor = p
        break
if anchor is None:
    # fallback to H2 "TÓM TẮT"
    for p in doc.paragraphs:
        if p.text.strip().startswith('TÓM TẮT (Abstract)'):
            anchor = p
            break

# Collect missing image files
missing = [
    ('p001_img1_Image21.jpg', 'Logo nhóm NAMHS quốc tế — Institutional logo', 'NAMHS international project logo'),
    ('p001_img2_Image24.png', 'Logo IOS / VASS — Institutional logo', 'Institute of Sociology (VASS) logo'),
    ('p001_img3_Image26.jpg', 'Logo UQ — University of Queensland', 'University of Queensland logo'),
    ('p001_img5_Image23.jpg', 'Logo JHSPH — Johns Hopkins Bloomberg', 'Johns Hopkins Bloomberg School of Public Health logo'),
]

# Create a "logos row" section before anchor
# Strategy: add logos in a row with small width
logo_para = doc.add_paragraph()
logo_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
title_run = logo_para.add_run('\nĐƠN VỊ CHỦ TRÌ (Implementing Institutions)')
title_run.font.name = 'Times New Roman'
title_run.font.size = Pt(11)
title_run.bold = True
title_run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x68)

for fname, cap_vn, cap_en in missing:
    path = os.path.join(IMG_DIR, fname)
    if not os.path.exists(path):
        print(f'  [SKIP] {fname}')
        continue
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(path, width=Cm(5.0))
    except Exception as e:
        print(f'  [ERR] {fname}: {e}')
        continue
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{cap_vn}\n({cap_en})')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    print(f'  OK: {fname}')

# Move the newly added logo block to right after "BẢNG VIẾT TẮT" table
# For simplicity, keep it at end of doc
doc.save(OUT)

# Verify
import zipfile
with zipfile.ZipFile(OUT) as z:
    n = len([x for x in z.namelist() if x.startswith('word/media/')])
print(f'\nFinal images in docx: {n}/12 (VNAMHS PDF has 12 images)')
