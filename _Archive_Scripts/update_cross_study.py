# -*- coding: utf-8 -*-
"""Cập nhật Cross-Study Synthesis — 35 bài NC"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '03_Ban-dich'))
from tao_dich_template import *

doc = create_doc()

add_heading(doc, 'TỔNG HỢP LIÊN BÀI BÁO\nLO ÂU Ở HỌC SINH THCS VÀ THPT', 1)
add_p(doc, 'Phân tích sâu từ 35 nghiên cứu (2020–2025)', bold=True)
add_p(doc, 'Việt Nam — Đông Nam Á — Thế giới | Cập nhật: 04/04/2026')
add_p(doc, '7 bài Việt Nam + 28 bài quốc tế | 35 bản dịch + 37 tóm tắt')

# ========== 1. TỶ LỆ LO ÂU THEO QUỐC GIA ==========
add_heading(doc, '1. Tỷ lệ lo âu theo quốc gia, công cụ và năm', 2)
add_p(doc, 'Bảng dưới đây tổng hợp tỷ lệ lo âu từ 35 NC tại 20+ quốc gia. Điểm quan trọng nhất: sự khác biệt rất lớn giữa các NC phản ánh khác biệt CÔNG CỤ ĐO, không nhất thiết khác biệt thực sự.')

add_heading(doc, 'Bảng 1. Tỷ lệ lo âu VTN — Tổng hợp 35 bài NC', 3)
add_table(doc,
    ['Quốc gia', 'Tỷ lệ', 'Công cụ', 'n', 'Nguồn'],
    [['VN — Hà Nội', '40,6%', 'GAD-7 (sàng lọc)', '3.910', 'Hoa 2024 (VN1)'],
     ['VN — Quốc gia', '2,3% (chẩn đoán)', 'DISC-5/DSM-5', '5.996', 'V-NAMHS 2022 (VN2)'],
     ['VN — DTTS Lạng Sơn', '54,4%', 'DASS-21', '845', 'Ngô Anh Vinh 2024 (VN15)'],
     ['VN — Hà Nội', '86,2%', 'DASS-21', '501', 'Bảo Quyên 2025 (VN16)'],
     ['VN — Thanh Hóa', '49,0%', 'DASS-21', '482', 'Danh Lâm 2022 (VN17)'],
     ['VN — An Giang', '61,2%', 'DASS-21', '366', 'An Giang 2025 (VN18)'],
     ['VN — Huế', '65,8%', 'DASS-21', '685', 'Thảo Vi 2025 (VN19)'],
     ['VN — TPHCM', '25,1%', 'DASS-Y (VTN riêng)', '976', 'Vĩnh Lộc 2024 (VN20)'],
     ['VN — 7 nước (SAD)', '30,7%', 'SIAS-17 (lo âu XH)', '6.825', 'Jefferies 2020 (QT35)'],
     ['Mỹ — hệ thống y tế', '9,6%→19,2%', 'Chẩn đoán lâm sàng', '13,7 triệu', 'JAACAP 2024 (QT23)'],
     ['Mỹ — trẻ em', 'Lo âu tăng 61%', 'NSCH', 'Quốc gia', 'NSCH 2020 (QT04)'],
     ['Châu Âu — WHO', '9 triệu VTN', 'Tổng hợp', '53 nước', 'WHO Europe 2025 (QT24)'],
     ['Anh', '6,3% (7–16 tuổi)', 'NHS Digital', 'Quốc gia', 'UK NHS 2024 (QT26)'],
     ['Na Uy', 'Tăng 13 năm', 'HSCL-6', '979.043', 'Norway 2025 (QT21)'],
     ['Ireland', 'Tăng 2012→2019', 'DASS-21', '11.954', 'Ireland 2024 (QT32)'],
     ['Hàn Quốc', 'Giảm→Tăng (COVID)', 'KYRBS', 'Quốc gia', 'Korea 2024 (QT34)'],
     ['Úc', 'K10 cao: 30,7%→40,7%', 'K10 + MHC-SF', '5.656', 'EpiPsychSci 2025 (QT25)'],
     ['Toàn cầu', 'AAPC 0,84%/năm', 'GBD 2021', '204 nước', 'GBD 2025 (QT30)'],
     ['59 nước LMIC', 'Phổ biến', 'GSHS', '179.937', '59 Countries 2025 (QT31)']],
    widths=[3.0, 3.0, 3.0, 2.0, 3.0])

add_p(doc, 'Phân tích:', bold=True)
add_p(doc, 'a) Tỷ lệ dao động từ 2,3% (chẩn đoán DISC-5) đến 86,2% (sàng lọc DASS-21) — khoảng cách gấp 37 lần. Nguyên nhân CHÍNH: khác biệt công cụ đo, KHÔNG phải khác biệt lo âu thực sự.')
add_p(doc, 'b) DASS-Y (VTN riêng) cho tỷ lệ 25,1% — THẤP HƠN đáng kể so với DASS-21 (40–86%). Gợi ý: DASS-21 người lớn có thể ĐÁNH GIÁ QUÁ MỨC lo âu ở VTN do ngưỡng không phù hợp.')
add_p(doc, 'c) Lo âu XÃ HỘI (SIAS): 30,7% VTN VN — loại lo âu riêng, không đo bởi GAD-7/DASS-21. 1/6 có SAD nhưng không biết ("lo âu ẩn").')

# ========== 2. XU HƯỚNG TOÀN CẦU ==========
add_heading(doc, '2. Xu hướng lo âu VTN toàn cầu — Bằng chứng mạnh', 2)

add_heading(doc, 'Bảng 2. Xu hướng lo âu VTN — 7 NC dài hạn', 3)
add_table(doc,
    ['Nước/Vùng', 'Giai đoạn', 'Xu hướng', 'Nguồn', 'Đặc thù'],
    [['Toàn cầu (204 nước)', '1990–2021 (31 năm)', 'AAPC 0,84%', 'GBD 2025 (QT30)', 'Tăng tốc từ 2014'],
     ['Hoa Kỳ', '2013–2021 (8 năm)', 'Lo âu TĂNG GẤP ĐÔI (AOR=2,17)', 'JAACAP 2024 (QT23)', 'Chẩn đoán lâm sàng'],
     ['Na Uy', '2011–2024 (13 năm)', 'Tăng liên tục', 'Norway 2025 (QT21)', 'Trường + MXH giải thích'],
     ['Hàn Quốc', '2006–2022 (16 năm)', 'Giảm trước COVID → Tăng sau', 'Korea 2024 (QT34)', 'ĐẢO CHIỀU — duy nhất'],
     ['Ireland', '2012–2019 (7 năm)', 'Tăng, nữ nhanh hơn', 'Ireland 2024 (QT32)', 'Trước COVID'],
     ['Úc', '2018–2022 (4 năm)', 'Flourishing giảm 53%→44%', 'EpiPsychSci 2025 (QT25)', 'CMH toàn diện'],
     ['Anh', '2017–2023', 'Lo âu 3,5%→6,3% (gấp đôi)', 'UK NHS 2024 (QT26)', 'Trẻ 7–16 tuổi']],
    widths=[3.0, 3.0, 3.5, 3.0, 3.0])

add_p(doc, 'Kết luận: Bằng chứng NHẤT QUÁN từ 7 NC dài hạn — lo âu VTN TĂNG trên toàn cầu. Hàn Quốc là ngoại lệ duy nhất (cải thiện trước COVID rồi xấu đi). VN THIẾU dữ liệu xu hướng dài hạn — GAP cấp thiết.', bold=True)

# ========== 3. NGUYÊN NHÂN XU HƯỚNG ==========
add_heading(doc, '3. Nguyên nhân xu hướng tăng — Bằng chứng liên bài', 2)

add_heading(doc, 'Bảng 3. Yếu tố giải thích xu hướng tăng lo âu VTN', 3)
add_table(doc,
    ['Yếu tố', 'Bằng chứng', 'Nguồn', 'Mức độ'],
    [['Bất mãn trường học', 'Giải thích CHÍNH (decomposition loại bỏ xu hướng)', 'Norway 2025 (QT21)', '★★★★★'],
     ['Mạng xã hội (thời gian)', 'VTN SKTT dùng nhiều hơn (g=0,46); giải thích một phần xu hướng', 'Nature 2025 (QT27); Norway (QT21)', '★★★★'],
     ['Mạng xã hội (tác động)', 'Cắt ngang mạnh, DỌC YẾU; RCT giảm ST→cải thiện SKTT (d=0,53)', 'Li 2025 (QT22); JAMA 2024 (QT33)', '★★★'],
     ['Áp lực học tập', 'ESSA ≥59 → tăng lo âu; OR=11,6 (áp lực→lo âu)', 'Vĩnh Lộc 2024 (VN20); Wen 2020', '★★★★'],
     ['Cấu trúc gia đình', 'Ly hôn, bỏ mặc, ACEs → tăng lo âu', 'VN20; Ngô Anh Vinh (VN15); Qiu 2022', '★★★★'],
     ['Bất an thực phẩm', 'AOR=2,22 (59 nước)', '59 Countries (QT31)', '★★★'],
     ['Bắt nạt', 'AOR=1,68 (59 nước)', 'QT31; Nakie 2022', '★★★'],
     ['Giấc ngủ', 'AOR=13,71 (<5h → trầm cảm chắc chắn)', 'Zhu 2025', '★★★★'],
     ['COVID-19', 'Đảo chiều Hàn Quốc; tăng tốc GBD từ 2019', 'Korea (QT34); GBD (QT30)', '★★★'],
     ['Bất bình đẳng thu nhập', 'Stress nghèo 62,8% vs giàu 40,1% (chênh 22,7)', 'Korea 2024 (QT34)', '★★★']],
    widths=[3.5, 5.0, 3.0, 2.0])

# ========== 4. GIỚI TÍNH ==========
add_heading(doc, '4. Giới tính — Nhất quán xuyên suốt (trừ lo âu XH)', 2)

add_heading(doc, 'Bảng 4. Khác biệt giới tính trong lo âu VTN', 3)
add_table(doc,
    ['Nghiên cứu', 'Nữ', 'Nam', 'Chỉ số', 'Ghi chú'],
    [['Hoa 2024 (VN)', 'M=1,74', 'M=1,50', 'p<0,01', 'Nữ cao hơn 16%'],
     ['Vĩnh Lộc 2024 (VN)', 'Cao hơn', 'Thấp hơn', 'p<0,05', 'Cả 3 chỉ số DAS'],
     ['JAACAP 2024 (Mỹ)', '22,3%', '16,5%', 'AOR=2,24 vs 2,14', 'Nữ rõ rệt hơn'],
     ['Ireland 2024', 'Tăng nhanh hơn', '—', 'p<0,01', 'Khoảng cách MỞ RỘNG'],
     ['Úc 2025 (CMH)', 'Flourishing 36,3%', '52,2%', '15,9 điểm chênh', 'Nữ xấu hơn nhiều'],
     ['59 Countries', 'AOR=1,51', 'Reference', 'p<0,001', 'Nhất quán 59 nước'],
     ['Lo âu XH 7 nước', '36,5%', '35,6%', 'n.s.', 'KHÔNG khác biệt — ngoại lệ']],
    widths=[3.0, 2.5, 2.5, 2.5, 3.5])

add_p(doc, 'Nữ > nam NHẤT QUÁN trong lo âu tổng quát — nhưng lo âu XÃ HỘI không khác biệt giới. Gợi ý: lo âu XH ảnh hưởng nam=nữ, khác cơ chế với GAD.')

# ========== 5. CAN THIỆP ==========
add_heading(doc, '5. Can thiệp — Bằng chứng từ 5 NC', 2)

add_heading(doc, 'Bảng 5. Hiệu quả can thiệp cho lo âu VTN', 3)
add_table(doc,
    ['Can thiệp', 'Hiệu quả', 'Nguồn', 'Áp dụng VN'],
    [['CBT + SSRI (kết hợp)', '80,7% đáp ứng (CAMS)', 'AJP 2024 (QT28)', 'Thiếu chuyên gia VN'],
     ['CBT đơn', '59,7% đáp ứng; NNT ≈ 3', 'AJP 2024 (QT28)', 'CBT nhóm khả thi'],
     ['ACT (Chấp nhận & Cam kết)', 'SUCRA 0,69 — hạng 1 NMA', 'BMC Psych 2025 (QT29)', 'Mới ở VN'],
     ['Hoạt động thể chất', 'Hiệu quả (SUCRA 0,51)', 'BMC Psych 2025 (QT29)', 'DỄ triển khai tại trường'],
     ['VRET (Thực tế ảo)', 'Hiệu quả (SUCRA 0,51)', 'BMC Psych 2025 (QT29)', 'Cần công nghệ'],
     ['Giảm screen time (RCT)', 'Cohen d = 0,53; nội hóa cải thiện', 'JAMA 2024 (QT33)', 'Can thiệp gia đình'],
     ['Trường học LMIC', '3/4 hiệu quả trầm cảm; 1/4 lo âu', 'Zhameden 2025', 'CHƯA CÓ RCT VN'],
     ['Hỗ trợ SKTT trường', 'OR = 0,562 (bảo vệ)', 'Wen 2020', 'Phù hợp VN'],
     ['Tăng lạc quan (positive psych)', 'β gián tiếp = −0,24 qua lo âu', 'Thảo Vi 2025 (VN19)', 'Khả thi']],
    widths=[4.0, 3.5, 3.0, 3.0])

# ========== 6. SCREEN TIME & MXH ==========
add_heading(doc, '6. Screen time & MXH — Tổng hợp 5 NC', 2)

add_heading(doc, 'Bảng 6. Bằng chứng screen time → SKTT VTN', 3)
add_table(doc,
    ['Nghiên cứu', 'Thiết kế', 'Phát hiện', 'Kết luận'],
    [['Norway 2025 (QT21)', 'Decomposition 13 năm', 'MXH giải thích MỘT PHẦN xu hướng', 'Đóng góp, không chính'],
     ['Li 2025 (QT22)', 'Dọc 12 tháng', 'Cắt ngang mạnh, DỌC YẾU (lo âu: p=0,443)', 'Quan hệ HAI CHIỀU?'],
     ['Nature 2025 (QT27)', 'Cắt ngang + chẩn đoán', 'VTN SKTT dùng MXH nhiều hơn (g=0,46)', 'Nội hóa > ngoại hóa'],
     ['JAMA 2024 (QT33)', 'RCT 2 tuần', 'Giảm ST → cải thiện SKTT (d=0,53)', 'NHÂN QUẢ — hiếm'],
     ['Hoàng Trung Học (VN14)', 'Cắt ngang VN', 'Game điện tử Beta = 0,176', 'Liên quan ở VN']],
    widths=[3.0, 2.5, 4.5, 3.0])

add_p(doc, 'Tổng hợp: Tác động screen time/MXH lên SKTT VTN TỒN TẠI nhưng PHỨC TẠP hơn giả định. Cắt ngang mạnh, dọc yếu — gợi ý quan hệ hai chiều. RCT (JAMA) cho thấy CẦN GIẢM MẠNH (>80%) mới hiệu quả. VTN nội hóa (lo âu) nhạy cảm hơn ngoại hóa.', bold=True)

# ========== 7. KHOẢNG CÁCH DỊCH VỤ ==========
add_heading(doc, '7. Khoảng cách dịch vụ SKTT', 2)

add_heading(doc, 'Bảng 7. Tiếp cận dịch vụ SKTT — So sánh quốc tế', 3)
add_table(doc,
    ['Nước', 'Tiếp cận dịch vụ', 'Chi tiêu', 'Nguồn'],
    [['VN', '8,4% VTN', 'Chưa có số liệu', 'V-NAMHS 2022'],
     ['Anh', 'NHS phổ quát', '£16,3 tỷ/năm (10,6% NHS)', 'UK NHS 2024 (QT26)'],
     ['Mỹ', '82,6% (ước tính)', 'Hệ thống phát triển', 'JAACAP 2024'],
     ['Philippines', '2%', 'Rất thấp', 'Puyat 2025'],
     ['Châu Âu', '16/27 nước thiếu', 'Khác nhau', 'WHO Europe 2025 (QT24)'],
     ['VN — phụ huynh', 'Chỉ 5,1% nhận ra', '—', 'V-NAMHS 2022']],
    widths=[2.5, 3.5, 3.5, 3.5])

# ========== 8. KHOẢNG TRỐNG NC ==========
add_heading(doc, '8. Khoảng trống nghiên cứu — Ưu tiên cho Việt Nam', 2)

add_heading(doc, 'Bảng 8. Top 10 khoảng trống NC — cập nhật 35 bài', 3)
add_table(doc,
    ['#', 'Khoảng trống', 'Bằng chứng', 'Ưu tiên'],
    [['1', 'RCT can thiệp SKTT tại trường VN', 'Zhameden 2025: 0 RCT từ VN; BMC NMA: CBT+PE hiệu quả', 'RẤT CAO'],
     ['2', 'So sánh sàng lọc vs chẩn đoán trên cùng mẫu VN', 'Hoa 40,6% (GAD-7) vs V-NAMHS 2,3% (DISC-5) — chênh 17 lần', 'RẤT CAO'],
     ['3', 'Khảo sát quốc gia lặp lại (xu hướng dài hạn VN)', 'GBD AAPC 0,84%; Norway/Mỹ/Hàn Quốc đều có xu hướng; VN chưa có', 'RẤT CAO'],
     ['4', 'Phân tích SKTT VTN VN theo thu nhập', 'Korea 2024: stress nghèo 62,8% vs giàu 40,1%; VN: chưa có dữ liệu', 'CAO'],
     ['5', 'NC lo âu XÃ HỘI chuyên biệt ở VTN VN', 'Jefferies 2020: VN SAD=30,7% nhưng 18% không biết; chưa có NC riêng', 'CAO'],
     ['6', 'Decomposition nguyên nhân xu hướng tại VN', 'Norway 2025: bất mãn trường + MXH; VN có dữ liệu nhưng chưa phân tích', 'CAO'],
     ['7', 'RCT giảm screen time tại trường VN', 'JAMA 2024: d=0,53 chỉ 2 tuần; Li 2025 dọc yếu; VN: 0 RCT', 'CAO'],
     ['8', 'Đánh giá DASS-Y vs DASS-21 vs GAD-7 ở VTN VN', 'Vĩnh Lộc: DASS-Y cho 25,1% vs DASS-21 40-86%; ngưỡng nào phù hợp?', 'TRUNG BÌNH'],
     ['9', 'NC SKTT VTN DTTS/vùng khó khăn', 'Ngô Anh Vinh: DTTS 54,4%; 59 Countries: thực phẩm AOR=2,22', 'CAO'],
     ['10', 'Bất an thực phẩm → lo âu VTN VN', '59 Countries (QT31): AOR=2,22 nhưng VN chưa NC; nông thôn/DTTS', 'TRUNG BÌNH']],
    widths=[0.8, 5.0, 5.5, 2.0])

# ========== 9. ĐỀ CƯƠNG GỢI Ý ==========
add_heading(doc, '9. Gợi ý đề cương nghiên cứu — 3 giai đoạn', 2)

add_heading(doc, 'Bảng 9. Đề cương gợi ý — Dựa trên 35 bài NC', 3)
add_table(doc,
    ['Giai đoạn', 'Nội dung', 'Phương pháp', 'Thời gian'],
    [['1. Khảo sát', 'Tỷ lệ lo âu (GAD-7 + DASS-Y + SIAS) + yếu tố nguy cơ/bảo vệ', 'Cắt ngang, 1.000+ HS THCS/THPT, 3 vùng', '6 tháng'],
     ['2. So sánh công cụ', 'GAD-7 vs DASS-21 vs DASS-Y vs SIAS trên cùng mẫu + DISC-5 mẫu con', 'Thiết kế so sánh công cụ, 300–500 HS', '6 tháng'],
     ['3. Can thiệp', 'RCT CBT nhóm + PE tại trường vs đối chứng', 'RCT cluster (5 trường can thiệp, 5 đối chứng)', '12 tháng']],
    widths=[2.0, 5.5, 4.0, 2.0])

add_p(doc, 'Giai đoạn 1 ưu tiên dùng: GAD-7 (so sánh Hoa 2024), DASS-Y (so sánh Vĩnh Lộc 2024), và SIAS-17 (so sánh Jefferies 2020 — có dữ liệu VN). Thêm ESSA (áp lực học tập), ACEs, screen time, giấc ngủ.')
add_p(doc, 'Giai đoạn 3: CBT nhóm + hoạt động thể chất (BMC NMA 2025: cả hai hiệu quả). Kết hợp giảm screen time (JAMA 2024). Đo trước/sau bằng GAD-7 + DASS-Y + SDQ.')

add_p(doc, '')
add_p(doc, 'Tài liệu cập nhật ngày 04/04/2026. Dựa trên 35 bài NC đã dịch và kiểm tra (2020–2025).', italic=True)

# SAVE
doc.save('Tổng hợp liên bài báo - Lo âu HS - 04042026.docx')

import docx as dx
d = dx.Document('Tổng hợp liên bài báo - Lo âu HS - 04042026.docx')
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Cross-study: {len(t)} chars, {len(d.tables)} tables')
print('DONE!')
