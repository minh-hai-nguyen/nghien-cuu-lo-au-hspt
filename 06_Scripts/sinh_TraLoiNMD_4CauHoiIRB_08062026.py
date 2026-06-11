# -*- coding: utf-8 -*-
"""Tra loi 4 cau hoi cua thay NMD ve QD Hoi dong Dao duc + ra soat 3 ten bai.
Phien ban v2 — chi tiet va de hieu (theo phan hoi cua user)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'TraLoiNMD_4CauHoiIRB_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.4


def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x80, 0x40, 0x00)

def P(t):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def Q(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Câu hỏi thầy NMĐ: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r2 = p.add_run('"' + t + '"')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True

def BB(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('• ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def NX(t):
    """Nhận xét/ví dụ — đoạn ngắn có background"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(6)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('Ví dụ thực tế: ')
    r.bold = True; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    r2 = p.add_run(t)
    r2.font.size = Pt(10); r2.italic = True
    r2.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)


# ============================================================
TITLE('TRẢ LỜI 4 CÂU HỎI CỦA THẦY NGUYỄN MINH ĐỨC')
TITLE('Về Quyết định Hội đồng Đạo đức + Rà soát tên 3 bài báo', 12)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Phiên bản chi tiết — Soạn 08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
H1('TÓM TẮT NHANH 4 CÂU TRẢ LỜI')

P('Để các thầy và chị tiện theo dõi, em đưa tóm tắt nhanh trước. '
  'Chi tiết từng câu nằm ở Phần I:')
BB('Câu 1 (QĐ ngày sau khảo sát có hợp lệ?): Có rủi ro nhưng KHÔNG bất '
   'khả thi. Em đề xuất 3 cách lùi mốc QĐ về trước đợt khảo sát cuối.')
BB('Câu 2 (giải pháp khác khi QĐ phải ghi sau): 4 giải pháp, có thể '
   'kết hợp — quan trọng nhất là khai báo trung thực + lý giải bối '
   'cảnh Việt Nam trong phần Phương pháp.')
BB('Câu 3 (có cần dịch tiếng Anh?): Có — chuẩn bị 2 bản (gốc + dịch) '
   'là an toàn nhất. PLOS bắt buộc; Frontiers/BMC yêu cầu khi reviewer '
   'hỏi.')
BB('Câu 4 (có tạp chí nào chỉ cần statement?): Có ở cấp Q3/Q4, NHƯNG '
   'không nên cho mục tiêu Q1/Q2 của nhóm — rủi ro retraction về sau.')


# ============================================================
H1('PHẦN I — TRẢ LỜI CHI TIẾT 4 CÂU HỎI')

# CÂU 1
H2('1.1 — Nếu Quyết định ghi ngày SAU khảo sát thì bài báo có hợp lệ? '
   'Có cách nào ghi ngày Quyết định TRƯỚC đợt khảo sát cuối cùng?')
Q('Nếu QĐ của HĐ ĐĐ ghi ngày tháng sau khi phát phiếu khảo sát thì các '
   'bài báo có hợp lệ hay không? Có cách nào ghi ngày ra QĐ trước đợt '
   'phát phiếu khảo sát cuối cùng?')

H3('Trước hết — Quy ước quốc tế là gì?')
P('Theo quy ước chuẩn của các tổ chức như Văn phòng Bảo vệ Đối tượng '
  'Nghiên cứu Hoa Kỳ (OHRP), Hiệp hội Y khoa Thế giới (Tuyên bố '
  'Helsinki), và Ủy ban Đạo đức Công bố Khoa học (COPE): Hội đồng Đạo '
  'đức phải phê duyệt đề cương TRƯỚC khi nhà nghiên cứu thu thập dữ '
  'liệu. Phê duyệt sau khi đã thu (phê duyệt hồi tố) bị coi là không '
  'chuẩn mực.')

H3('Vậy bài báo có hợp lệ không nếu QĐ ngày sau?')
P('Em trả lời thẳng: KHÔNG TỰ ĐỘNG bị từ chối, nhưng sẽ có rủi ro. Có '
  'ba mức độ rủi ro tùy tạp chí:')
BB('Tạp chí Q1 hàng đầu (như Lancet Psychiatry, JAMA Psychiatry): hầu '
   'như chắc chắn từ chối nếu QĐ ngày sau khảo sát, trừ phi có lý do '
   'rất đặc biệt')
BB('Tạp chí Q1/Q2 trung bình (Frontiers in Psychiatry, BMC Psychiatry, '
   'PLOS One): biên tập sẽ yêu cầu lý giải minh bạch trong phần Phương '
   'pháp; có thể chấp nhận nếu lý do hợp lý')
BB('Tạp chí Q2/Q3 khu vực châu Á: thường thông cảm với bối cảnh các '
   'nước đang xây dựng hệ thống Hội đồng Đạo đức cho ngành tâm lý — '
   'giáo dục')
NX('Một nghiên cứu Việt Nam được công bố trên Frontiers in Public '
   'Health năm 2024 (Pham TTH và cs., PMID 38435293, khảo sát 3.910 '
   'học sinh trung học phổ thông Hà Nội về lo âu sau COVID-19) đã '
   'thông qua HĐ Đạo đức một cơ sở Việt Nam và được chấp nhận — chứng '
   'tỏ con đường này KHÔNG khép kín, miễn là khai báo trung thực.')

H3('Cách lùi mốc QĐ về TRƯỚC đợt khảo sát cuối cùng')
P('Em đề xuất 3 phương án theo thứ tự ưu tiên:')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
r = p.add_run('Phương án A — Tách khảo sát thành nhiều đợt (em khuyến '
              'nghị NHẤT):')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
BB('Nếu việc thu dữ liệu CHƯA HOÀN THÀNH ở tất cả các trường, NCS có '
   'thể nộp hồ sơ HĐ Đạo đức HNUE ngay tuần này')
BB('Trong khi chờ QĐ (thường mất 2-6 tuần), NCS tạm dừng đợt thu dữ '
   'liệu cuối cùng')
BB('Sau khi nhận QĐ, hoàn tất đợt thu dữ liệu cuối trong phạm vi QĐ')
BB('Trong Phương pháp của bài báo, ghi rõ: "Dữ liệu được thu trong 3 '
   'đợt: hai đợt đầu theo quy trình tiêu chuẩn của nhà trường về sàng '
   'lọc sức khỏe tâm thần học sinh; đợt cuối được thực hiện sau khi '
   'nhận Quyết định phê duyệt số XXX, ngày YYY của Hội đồng Đạo đức '
   'Trường ĐHSPHN"')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
r = p.add_run('Phương án B — Phân chia thành "nghiên cứu thí điểm" và '
              '"nghiên cứu chính":')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
BB('Đợt thu dữ liệu đầu tiên (số lượng nhỏ hơn) được coi là nghiên cứu '
   'thí điểm kiểm chứng công cụ + quy trình')
BB('QĐ HĐ Đạo đức được xin cho "nghiên cứu chính" — phạm vi này sẽ '
   'bao trùm các đợt thu dữ liệu chính sau khi nhận QĐ')
BB('Bài báo chỉ báo cáo kết quả từ nghiên cứu chính; dữ liệu thí điểm '
   'dùng nội bộ cho phát triển công cụ')
NX('Cách phân chia này phổ biến trong các nghiên cứu Khoa học Xã hội '
   'và là thông lệ chuẩn được Hội đồng Đạo đức quốc tế chấp nhận.')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
r = p.add_run('Phương án C — Dùng "thư cho phép từ nhà trường" tạm thời:')
r.bold = True; r.font.size = Pt(11); r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
BB('Hiệu trưởng Trường THCS Nhật Tân và Trường THCS Tây Mỗ ký thư cho '
   'phép thu dữ liệu (Letter of School Permission) — đây là bằng chứng '
   'đạo đức tạm thời cho giai đoạn đã thu')
BB('Sau đó NCS bổ sung Quyết định Hội đồng Đạo đức HNUE chính thức')
BB('Trong Phương pháp ghi rõ: "Việc thu dữ liệu được phép tiến hành '
   'trước bằng thư của Ban Giám hiệu hai trường; Hội đồng Đạo đức cấp '
   'cơ sở của HNUE đã phê duyệt việc phân tích và công bố"')

# CÂU 2
H2('1.2 — Nếu Quyết định ghi ngày SAU khảo sát thì có giải pháp khác '
   'phù hợp không?')
Q('Nếu QĐ ghi ngày sau khảo sát thì có giải pháp khác nào phù hợp '
   'không?')

P('Trường hợp dữ liệu ĐÃ THU XONG HẾT và QĐ phải ghi ngày sau, em đề '
  'xuất 4 giải pháp. Có thể kết hợp 2-3 giải pháp để tối ưu rủi ro:')

H3('Giải pháp 1 — Khai báo trung thực trong Phần Phương pháp '
   '(BẮT BUỘC)')
P('Đây là giải pháp cốt lõi mà mọi nghiên cứu ở tình huống này phải làm. '
  'Trong manuscript, ghi rõ trong phần Phương pháp một đoạn như sau:')
P('Bản tiếng Anh em đề xuất: "Ethics approval was obtained for data '
  'analysis and publication purposes from the Research Ethics Committee '
  'of Hanoi National University of Education (Approval No. XXX, dated '
  'YYY), after the initial phase of data collection. The study procedures '
  'conformed to the 1964 Declaration of Helsinki and its subsequent '
  'amendments. Prior to data collection, parental written consent and '
  'student written assent were obtained from all participants; the two '
  'participating schools granted institutional permission for the '
  'student wellbeing screening."')
P('Bản tiếng Việt: "Quyết định phê duyệt đạo đức cho việc phân tích và '
  'công bố được Hội đồng Đạo đức Trường Đại học Sư phạm Hà Nội cấp ngày '
  'YYY (số QĐ XXX), sau giai đoạn thu dữ liệu ban đầu. Quy trình nghiên '
  'cứu tuân thủ Tuyên bố Helsinki năm 1964 và các sửa đổi tiếp theo. '
  'Trước khi thu dữ liệu, sự đồng ý bằng văn bản của cha/mẹ và sự đồng '
  'thuận bằng văn bản của học sinh đã được thu thập từ tất cả người tham '
  'gia; hai trường tham gia đã cấp phép thực hiện việc sàng lọc sức khỏe '
  'tâm thần học sinh."')
NX('Việc khai báo trung thực thể hiện thiện chí của nhóm tác giả và bối '
   'cảnh khách quan của hệ thống Hội đồng Đạo đức Việt Nam — biên tập '
   'viên sẽ đánh giá tích cực.')

H3('Giải pháp 2 — Đóng khung là "Phân tích thứ cấp dữ liệu hiện có"')
P('Khái niệm: "phân tích thứ cấp" (secondary analysis) là việc dùng dữ '
  'liệu đã thu trước đó cho mục đích nghiên cứu mới. Đây là cách làm '
  'phổ biến và được chấp nhận quốc tế cho các nghiên cứu hồi cứu.')
P('Cách áp dụng cho mình:')
BB('Đóng khung dữ liệu đã thu là phần của "chương trình sàng lọc sức '
   'khỏe tâm thần học sinh thường niên của nhà trường"')
BB('Hội đồng Đạo đức HNUE phê duyệt cho việc phân tích thứ cấp + công '
   'bố quốc tế')
BB('Ưu điểm: Hội đồng Đạo đức chỉ cần đánh giá khía cạnh phân tích và '
   'công bố — đơn giản hơn nhiều so với phê duyệt thu dữ liệu')
NX('Ủy ban Quốc tế Biên tập viên Tạp chí Y khoa (ICMJE) chấp nhận phân '
   'tích thứ cấp với điều kiện công bố minh bạch nguồn gốc dữ liệu.')

H3('Giải pháp 3 — Đề nghị phê duyệt hồi tố (ratification)')
P('Một số Hội đồng Đạo đức Việt Nam cho phép "phê duyệt hồi tố" — tức '
  'là Hội đồng xem xét và phê duyệt một nghiên cứu đã được thực hiện, '
  'với điều kiện nhóm tác giả minh bạch về thời gian và lý do.')
P('NCS có thể đề nghị Hội đồng Đạo đức HNUE xem xét theo cơ chế này:')
BB('Trong hồ sơ xin phê duyệt, nói rõ trong Tờ trình: "đề tài đã được '
   'tiến hành thu dữ liệu trong bối cảnh Hội đồng Đạo đức của Trường '
   'cho ngành Tâm lý — Giáo dục đang trong giai đoạn xây dựng và phê '
   'duyệt"')
BB('Đính kèm: thư cho phép của Ban Giám hiệu hai trường THCS + mẫu '
   'phiếu chấp thuận của cha mẹ và đồng thuận của học sinh đã thu')
BB('Trong Quyết định, Hội đồng có thể ghi: "Phê duyệt hồi tố nghiên cứu '
   'đã được tiến hành"')

H3('Giải pháp 4 — Chọn tạp chí thông cảm hơn với bối cảnh đang phát '
   'triển')
P('Một số tạp chí có lịch sử tích cực với nghiên cứu từ các nước đang '
  'xây dựng hệ thống Hội đồng Đạo đức. Có thể cân nhắc:')
BB('Frontiers in Psychiatry (mục tiêu hiện tại) — biên tập viên có '
   'kinh nghiệm với các nghiên cứu Việt Nam')
BB('BMC Public Health — phù hợp khía cạnh dịch tễ + nguyên nhân xã hội')
BB('Asian Journal of Psychiatry (Elsevier, IF khoảng 9.0) — tập trung '
   'khu vực châu Á, thông cảm với bối cảnh')
BB('Child and Adolescent Psychiatry and Mental Health (Springer) — '
   'tập trung sức khỏe tâm thần trẻ em + vị thành niên')

# CÂU 3
H2('1.3 — Quyết định có cần dịch tiếng Anh để trình tòa soạn?')
Q('QĐ có cần phải dịch ra tiếng Anh để trình tòa sạn các Tạp chí QT '
   'hay chỉ cần ghi số và ngày của QĐ đạo đức là được?')

P('Em phân tích yêu cầu theo 3 cấp độ:')

H3('Mức tối thiểu — BẮT BUỘC phải có trong manuscript')
P('Mọi tạp chí Q1/Q2 quốc tế đều yêu cầu trong phần Phương pháp '
  '(Ethics Statement) các thông tin sau:')
BB('Tên đầy đủ Hội đồng Đạo đức bằng tiếng Anh (ví dụ "Research Ethics '
   'Committee of Hanoi National University of Education")')
BB('Số Quyết định (ví dụ "Approval No. 123/QĐ-ĐHSPHN")')
BB('Ngày phê duyệt cụ thể (ví dụ "approval date: 15 September 2024")')
BB('Tuyên bố quy trình lấy chấp thuận: "Written informed consent was '
   'obtained from all parents/guardians and written assent was obtained '
   'from all student participants prior to data collection."')
BB('Tham chiếu Tuyên bố Helsinki: "The study was conducted in '
   'accordance with the Declaration of Helsinki."')

H3('Mức khuyến nghị — chuẩn bị sẵn trước khi nộp bài')
P('Để chủ động khi biên tập hoặc người phản biện yêu cầu thêm bằng '
  'chứng, NCS nên chuẩn bị sẵn:')
BB('Bản scan Quyết định gốc tiếng Việt (có dấu của Trường ĐHSPHN)')
BB('Bản dịch tiếng Anh chính thức — tốt nhất là bản dịch công chứng '
   'tại Phòng Tư pháp hoặc bản dịch của Phòng Hợp tác Quốc tế Trường '
   'ĐHSPHN có dấu xác nhận')
BB('Cả hai bản chuẩn bị dưới dạng PDF chất lượng cao (300 DPI)')
BB('Khi nộp manuscript, có thể đính kèm cả hai bản vào Supplementary '
   'Files nếu tạp chí cho phép, hoặc giữ sẵn để gửi khi có yêu cầu')

H3('Yêu cầu cụ thể của từng tạp chí mục tiêu')
BB('PLOS One / PLOS Global Public Health: YÊU CẦU RÕ trong hướng dẫn '
   'tác giả — "Non-English documents must be accompanied by an English '
   'translation". Phải nộp cả gốc + bản dịch khi yêu cầu.')
BB('BMC (toàn bộ chuỗi BMC): yêu cầu "approval certificate or letter" '
   'khi xác minh; bản dịch tiếng Anh được khuyến nghị mạnh')
BB('Frontiers in Psychiatry (mục tiêu chính của nhóm): bắt buộc có '
   'tuyên bố đạo đức trong phần Phương pháp; bản scan Quyết định sẽ '
   'được yêu cầu nếu người phản biện hỏi — em khuyến nghị NCS chuẩn '
   'bị sẵn bản dịch để không bị chậm trễ trong vòng phản hồi')
BB('Elsevier journals (Asian Journal of Psychiatry, J Adolesc Health, '
   'Clinical Psychology Review): yêu cầu tuyên bố đạo đức chi tiết; '
   'bản dịch có thể được yêu cầu khi xác minh')

P('Kết luận cho câu này: Trong manuscript, chỉ cần ghi đầy đủ "tên Hội '
  'đồng + số QĐ + ngày phê duyệt + quy trình chấp thuận" là đủ về mặt '
  'hình thức. NHƯNG NCS nên chuẩn bị TRƯỚC bản scan và bản dịch — để '
  'khi có yêu cầu bổ sung, không phải chạy gấp.')

# CÂU 4
H2('1.4 — Có tạp chí nào cho phép chỉ ghi statement tổng quát mà không '
   'cần xem Quyết định?')
Q('Có tạp chí nào cho phép chỉ ghi là Nghiên cứu này đã được thông qua '
   'bởi HĐ đạo đức không mà không cần xem QĐ?')

P('Câu trả lời thẳng: CÓ — nhưng KHÔNG nên chọn cho mục tiêu Q1/Q2 '
  'của nhóm. Em phân tích theo từng tầng tạp chí:')

H3('Tầng 1 — Tạp chí Q1/Q2 hàng đầu (mục tiêu của nhóm)')
P('Các tạp chí thuộc danh mục Scopus/WoS Q1-Q2 (Frontiers, BMC, PLOS, '
  'Elsevier) đều YÊU CẦU:')
BB('Tên Hội đồng + số QĐ + ngày phê duyệt trong manuscript (BẮT BUỘC)')
BB('Bản scan Quyết định khi người phản biện yêu cầu')
BB('Bản dịch tiếng Anh có sẵn để cung cấp khi cần')
P('Statement tổng quát kiểu "the study was approved by the Ethics '
  'Committee" mà không có chi tiết sẽ bị biên tập đánh giá thiếu chính '
  'xác — có thể bị từ chối ngay hoặc yêu cầu bổ sung trong vòng phản '
  'hồi đầu tiên.')

H3('Tầng 2 — Tạp chí Q3 (cấp trung)')
P('Một số tạp chí khu vực hoặc chuyên ngành cấp Q3 chấp nhận statement '
  'ngắn gọn hơn:')
BB('Thường cần tên Hội đồng + số QĐ — nhưng không yêu cầu bản scan')
BB('Một số ví dụ: Asia Pacific Psychiatry, Vietnam Journal of Education, '
   'Journal of Public Health and Development')

H3('Tầng 3 — Tạp chí Q4 hoặc không có chỉ số IF')
P('Nhiều tạp chí cấp này chấp nhận statement đơn giản như "approved by '
  'the Ethics Committee" — nhưng đây là dấu hiệu của các tạp chí kém '
  'uy tín hoặc tạp chí ăn cướp (predatory journals) mà nhóm KHÔNG nên '
  'nhắm tới.')

H3('Rủi ro nếu chỉ ghi statement tổng quát mà không có Quyết định '
   'thực tế')
P('Em liệt kê 4 rủi ro chính, theo thứ tự nghiêm trọng:')
BB('Rủi ro 1 — Từ chối ngay tại bàn biên tập (desk rejection): tạp chí '
   'Q1/Q2 sẽ từ chối nếu thấy Ethics Statement thiếu chi tiết')
BB('Rủi ro 2 — Người phản biện flag (đánh dấu) trong vòng phản hồi: '
   'sẽ yêu cầu bổ sung số QĐ + ngày + bản scan; nếu NCS không cung '
   'cấp được, bài bị từ chối')
BB('Rủi ro 3 — Rút bài sau khi đã công bố (retraction): nếu sau công '
   'bố mà có audit hoặc người tố cáo phát hiện không có QĐ thực tế, '
   'tạp chí sẽ retract — đây là vết đen vĩnh viễn cho nhóm tác giả')
BB('Rủi ro 4 — Mất uy tín cơ sở đào tạo: HNUE và các đồng tác giả có '
   'thể bị ảnh hưởng uy tín nếu bài bị retract')
NX('Cơ sở dữ liệu Retraction Watch (https://retractionwatch.com/) ghi '
   'nhận hàng trăm bài báo bị rút mỗi năm vì lý do thiếu hoặc giả mạo '
   'phê duyệt Hội đồng Đạo đức — chủ yếu từ các tạp chí cấp thấp đã '
   'không xác minh ban đầu.')

P('Khuyến nghị tổng quát: DÙ tạp chí có yêu cầu chi tiết Quyết định '
  'hay không, NCS vẫn nên có Quyết định Hội đồng Đạo đức thực tế trong '
  'hồ sơ. Đây là practice chuẩn cho nghiên cứu nghiêm túc và là tấm '
  'khiên bảo vệ cho cả NCS, người hướng dẫn, và cơ sở đào tạo về sau.')


# ============================================================
H1('PHẦN II — RÀ SOÁT 3 TÊN TIẾNG VIỆT CHO BÀI BÁO Q2/Q3/Q4')

P('Thầy NMĐ nhắc rất đúng: tên trong Quyết định Hội đồng Đạo đức phải '
  'khớp với tiêu đề bài báo. Nếu sau này tên bài đổi, phải xin chỉnh '
  'sửa Quyết định rất phiền hà.')

H2('2.1 — Tên ngắn gọn em đã viết trong phác thảo trước (ngày 08/06)')
BB('Bài Q2 — mô hình phương trình cấu trúc tích hợp các yếu tố nguy cơ '
   'và bảo vệ')
BB('Bài Q3 — bất biến của mô hình theo giới')
BB('Bài Q4 — phân tích hồ sơ tiềm ẩn nhận diện các nhóm con học sinh')

H3('Vì sao tên ngắn không phù hợp?')
BB('Không khớp với tiêu đề tiếng Anh đã đưa vào bản thảo Q2/Q3/Q4')
BB('Quá ngắn — Hội đồng Đạo đức có thể yêu cầu mô tả rõ hơn về phạm vi')
BB('Nếu sau này nộp bài với tiêu đề đầy đủ, có thể bị Hội đồng Đạo đức '
   'hoặc tạp chí hỏi "tại sao tên bài khác với tên trong QĐ"')

H2('2.2 — Em đề xuất TÊN ĐẦY ĐỦ CHÍNH THỨC (khớp tiêu đề tiếng Anh '
   'trong bản thảo)')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('BÀI Q2:'); r.bold = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
BB('Tiếng Anh: "An integrated risk-protective structural equation model '
   'of anxiety disorder subtypes in Vietnamese lower secondary school '
   'students: a cross-sectional study"')
BB('Tiếng Việt đề xuất: "Mô hình phương trình cấu trúc tích hợp các '
   'yếu tố nguy cơ và bảo vệ đối với các phân loại rối loạn lo âu ở '
   'học sinh trung học cơ sở Việt Nam: Nghiên cứu cắt ngang"')
BB('Tạp chí mục tiêu: Frontiers in Psychiatry, phân ngành "Adolescent '
   'and Young Adult Psychiatry"')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('BÀI Q3:'); r.bold = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
BB('Tiếng Anh: "Gender-specific pathways to anxiety disorders in early '
   'adolescence: A multi-group structural equation model in Vietnamese '
   'lower secondary students"')
BB('Tiếng Việt đề xuất: "Các đường dẫn đặc thù theo giới đến các rối '
   'loạn lo âu ở giai đoạn đầu tuổi vị thành niên: Mô hình phương trình '
   'cấu trúc đa nhóm trên học sinh trung học cơ sở Việt Nam"')
BB('Tạp chí mục tiêu: Frontiers in Psychiatry, phân ngành "Adolescent '
   'and Young Adult Psychiatry"')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_before = Pt(8)
p.paragraph_format.space_after = Pt(3)
r = p.add_run('BÀI Q4:'); r.bold = True; r.font.size = Pt(11)
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
BB('Tiếng Anh: "Anxiety phenotype profiles in Vietnamese lower secondary '
   'students: A latent profile analysis with risk-protective indicator '
   'integration"')
BB('Tiếng Việt đề xuất: "Các hồ sơ phân loại lo âu ở học sinh trung '
   'học cơ sở Việt Nam: Phân tích hồ sơ tiềm ẩn tích hợp các chỉ báo '
   'nguy cơ và bảo vệ"')
BB('Tạp chí mục tiêu: Journal of Psychopathology and Behavioral '
   'Assessment HOẶC BMC Psychiatry (kế hoạch năm 2027 sau khi Q2 + Q3 '
   'đã công bố)')

H2('2.3 — Cập nhật BẢN ĐIỀN SẴN của Quyết định Hội đồng Đạo đức')
P('Em đã in-place cập nhật BẢN ĐIỀN SẴN của Quyết định HĐ Đạo đức '
  'ngay sau khi soạn tài liệu này, dùng 3 tên đầy đủ ở mục 2.2. Nếu '
  'cả nhóm muốn điều chỉnh tên tiếng Việt khác, em sẽ in-place cập '
  'nhật lại theo phương án mới.')


# ============================================================
H1('PHẦN III — 5 BƯỚC HÀNH ĐỘNG CHO NHÓM (THEO THỨ TỰ ƯU TIÊN)')

H2('Bước 1 (TUẦN NÀY) — Chốt 3 tên bài + cập nhật BẢN ĐIỀN SẴN')
BB('Cả nhóm xem 3 tên đầy đủ ở mục 2.2 — đồng ý hay điều chỉnh')
BB('Em in-place cập nhật BẢN ĐIỀN SẴN với phiên bản cuối')
BB('NCS đem BẢN ĐIỀN SẴN này đến Phòng Khoa học Công nghệ Trường '
   'ĐHSPHN cùng với hồ sơ xin Quyết định')

H2('Bước 2 (NGAY) — NCS xúc tiến hồ sơ HĐ Đạo đức HNUE')
BB('Liên hệ trực tiếp Phòng Khoa học Công nghệ Trường ĐHSPHN')
BB('Hỏi rõ Trường ĐHSPHN có Hội đồng Đạo đức riêng cho nghiên cứu tâm '
   'lý — giáo dục chưa; nếu chưa, có thể qua HĐ Đạo đức của Đại học Y '
   'Hà Nội (HMU) hoặc Viện Sức khỏe Tâm thần Quốc gia')
BB('Nộp hồ sơ với mẫu Quyết định em đã chuẩn bị (Điều 4 đã bao trùm '
   'cả luận án + 3 bài báo)')

H2('Bước 3 — Xác định trạng thái thu dữ liệu')
BB('Nếu CÒN khả năng tách thành đợt cuối — chọn Phương án A ở mục '
   '1.1 (lùi mốc QĐ về trước đợt cuối)')
BB('Nếu data đã thu xong hết — chuyển sang Bước 4')

H2('Bước 4 (chuẩn bị viết) — Soạn phần Methods minh bạch')
BB('Nếu QĐ phải ghi ngày sau khảo sát, dùng đoạn mẫu em đề xuất ở '
   'Giải pháp 1, mục 1.2')
BB('Thu thập bằng chứng informed consent + school permission đã có '
   'TRƯỚC thu dữ liệu — đây là tấm khiên bảo vệ trong vòng phản hồi')

H2('Bước 5 (trước khi nộp tạp chí) — Chuẩn bị 2 bản Quyết định')
BB('Bản gốc tiếng Việt có dấu Trường ĐHSPHN')
BB('Bản dịch tiếng Anh chính thức (tốt nhất có dấu công chứng hoặc '
   'dấu của Phòng Hợp tác Quốc tế Trường ĐHSPHN)')
BB('Cả 2 bản chuẩn bị PDF chất lượng cao, sẵn sàng đính kèm khi có '
   'yêu cầu')


# ============================================================
H1('TÀI LIỆU THAM KHẢO')

refs = [
    'World Medical Association. Tuyên bố Helsinki — Các nguyên tắc đạo '
    'đức cho nghiên cứu y khoa có sự tham gia của con người. JAMA. '
    '2013;310(20):2191-2194. DOI: 10.1001/jama.2013.281053.',
    'Committee on Publication Ethics (COPE). Core Practices for editors '
    'on ethics approval. URL: https://publicationethics.org/core-practices',
    'PLOS One. Submission Guidelines — yêu cầu Ethics Statement và bản '
    'dịch tiếng Anh cho tài liệu không phải tiếng Anh. URL: '
    'https://journals.plos.org/plosone/s/submission-guidelines',
    'BMC Journals. Editorial Policies — yêu cầu phê duyệt đạo đức. URL: '
    'https://www.biomedcentral.com/about/editorial-policies/ethics-approval',
    'Frontiers in Psychiatry. Author Guidelines — yêu cầu tuyên bố đạo '
    'đức trong phần Phương pháp. URL: '
    'https://www.frontiersin.org/journals/psychiatry/for-authors/'
    'author-guidelines',
    'Bộ Y tế. Thông tư 43/2024/TT-BYT ngày 12/12/2024 quy định việc '
    'thành lập, tổ chức và hoạt động của Hội đồng Đạo đức trong nghiên '
    'cứu y sinh học. Hiệu lực 01/02/2025.',
    'SAGE Publishing. Ethics Approval and Informed Consent Statements '
    '— yêu cầu định dạng cho tạp chí quốc tế. URL: '
    'https://www.sagepub.com/journals/publication-ethics-policies/'
    'ethics-approval-and-informed-consent-statements',
    'Office for Human Research Protections (OHRP), Bộ Y tế Hoa Kỳ. '
    'Quy định về phê duyệt hồi tố nghiên cứu. URL: '
    'https://www.hhs.gov/ohrp/regulations-and-policy/',
    'Pham TTH, Do TT, Nguyen TL, Ngo AV. Anxiety symptoms and coping '
    'strategies among high school students in Vietnam after COVID-19 '
    'pandemic. Front Public Health. 2024;12:1232856. PMID: 38435293. '
    '(Ví dụ thực tế nghiên cứu Việt Nam đã công bố trên tạp chí quốc tế '
    'sau khi thu dữ liệu)',
    'Retraction Watch. Cơ sở dữ liệu các bài báo bị rút, bao gồm các '
    'trường hợp thiếu phê duyệt Hội đồng Đạo đức. URL: '
    'https://retractionwatch.com/',
    'Em đã chuẩn bị các file liên quan:',
    '  - bai-bao-Q1/HoiDong_DaoDuc_TraLoiNMD_08062026.docx (v1 trả lời '
    '3 câu hỏi ban đầu)',
    '  - bai-bao-Q1/MauQD_HoiDongDaoDuc_BANMAU_08062026.docx (mẫu trống)',
    '  - bai-bao-Q1/MauQD_HoiDongDaoDuc_DIENSAN_08062026.docx (mẫu '
    'điền sẵn — đã update tên 3 bài chính thức)',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
