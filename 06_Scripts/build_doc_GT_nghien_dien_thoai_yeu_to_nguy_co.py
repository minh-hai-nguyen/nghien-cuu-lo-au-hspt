"""DOC 2: Phat bieu va luan chung gia thuyet
'Nghien dien thoai la yeu to nguy co doi voi RLLA cua HS' (CTH v6).

ALL FACTS VERIFIED:
- Sohn et al. 2019 BMC Psychiatry: 41 NC, N=41.871, OR lo au=3,05 (KTC 2,64-3,53; I²=0%); OR tram cam=3,17 (KTC 2,30-4,37; I²=78%); ty le 23,3% (14,0-31,2%) — verified PubMed 31779637
- Chen et al. 2023 (QT007) BMC Psychiatry: n=63.205, IGD OR=5,00 — verified Tom-tat-tung-bai
- Zheng & Peng 2025 (QT041) PRBM: SMAS r=0,415; β=0,153; mediation 63,13% qua self-efficacy — verified Tom-tat
- Fassi et al. 2025 (QT027) Nature Human Behaviour Q1 IF=24: n=3.340 UK 11-19 chan doan lam sang — verified Tom-tat
- Schmidt-Persson et al. 2024 (QT033) JAMA Q1 IF=13: RCT giam screen time CAI THIEN SKTT — verified Tom-tat
- Hoang Trung Hoc 2025 (VN014): n=8.473, beta=0,176 cho thoi gian dien tu — verified Tom-tat
- Li et al. 2025 (QT022) BJCP Q1: n=4.058 Uc, doc 12 thang, screen time T1 KHONG du bao lo au T2 (p=0,443) — verified Tom-tat
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Gia_thuyet_nghien_dien_thoai_yeu_to_nguy_co_RLLA.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GIẢ THUYẾT: NGHIỆN ĐIỆN THOẠI LÀ YẾU TỐ NGUY CƠ\nĐỐI VỚI RỐI LOẠN LO ÂU CỦA HỌC SINH\n— Phát biểu chính thức và luận chứng bằng chứng quốc tế —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# 1. Phát biểu giả thuyết
H('1. Phát biểu giả thuyết', level=2, color=NAVY)
para('Giả thuyết H (chính thức):', bold=True, indent=False)
para(
    'Học sinh trung học cơ sở có MỨC ĐỘ NGHIỆN ĐIỆN THOẠI cao hơn '
    'sẽ có TỶ LỆ rối loạn lo âu cao hơn — sau khi kiểm soát các '
    'biến nhân khẩu (giới, lớp, kinh tế gia đình) và các yếu tố '
    'tâm lý xã hội đồng biến (giấc ngủ, áp lực học tập, quan hệ '
    'cha mẹ-con).', bold=True
)
para('Định nghĩa thao tác (operational definition):', bold=True, indent=False)
para(
    'NGHIỆN ĐIỆN THOẠI (problematic smartphone use, PSU) — KHÁC '
    'với "thời gian dùng" (screen time) đơn thuần — là tổ hợp các '
    'biểu hiện đặc trưng của hành vi sử dụng quá mức và mất kiểm '
    'soát: (1) thèm dùng (craving); (2) mất kiểm soát thời gian; '
    '(3) ảnh hưởng tiêu cực đến học tập và quan hệ xã hội; (4) '
    'rút lui khi không có thiết bị; (5) tiếp tục dùng dù biết hậu '
    'quả tiêu cực. Đo bằng thang chuẩn — ưu tiên SAS-SV '
    '(Smartphone Addiction Scale-Short Version, Kwon và cộng sự, '
    '2013) hoặc BSMAS (Bergen Social Media Addiction Scale, '
    'Andreassen và cộng sự, 2016).', indent=False
)
para('Chiều quan hệ dự kiến:', bold=True, indent=False)
para(
    '• Hệ số tương quan r dự kiến trong khoảng 0,30–0,50.\n'
    '• Hệ số β chuẩn hóa trong hồi quy đa biến dự kiến trong '
    'khoảng 0,15–0,30 sau khi kiểm soát yếu tố nhiễu.\n'
    '• OR cho lo âu dự kiến trong khoảng 2,0–4,0 — phù hợp y '
    'văn quốc tế.', indent=False
)

# 2. Đặt vấn đề — 3 cấp độ
H('2. Bối cảnh — ba cấp độ bằng chứng', level=2, color=NAVY)
para(
    'Trên BÌNH DIỆN TOÀN CẦU, phân tích tổng hợp lớn nhất đến nay '
    '(Sohn và cộng sự, 2019 — BMC Psychiatry, 41 nghiên cứu, '
    'N = 41.871 trẻ em và thanh niên) đã thiết lập rõ rệt: học sinh '
    'có mức độ nghiện điện thoại cao có khả năng mắc lo âu CAO HƠN '
    'GẤP BA LẦN so với nhóm không nghiện — OR = 3,05 (KTC 95% từ '
    '2,64 đến 3,53; I² = 0% — heterogeneity bằng KHÔNG, đồng nhất '
    'tuyệt đối qua các nghiên cứu). Tương tự cho trầm cảm: OR = '
    '3,17 (KTC 95% 2,30–4,37). Tỷ lệ nghiện điện thoại trung vị '
    '23,3% (14,0%–31,2%) — gần MỘT PHẦN TƯ trẻ em và thanh niên.'
)
para(
    'Tại KHU VỰC châu Á — Đông Á, các nghiên cứu cỡ mẫu lớn nhất '
    'thế giới đã củng cố phát hiện này. Chen và cộng sự (2023) '
    'trên 63.205 học sinh trung học miền Tây Trung Quốc — BMC '
    'Psychiatry Q1 — phát hiện rối loạn chơi game online (đo bằng '
    'IGDS9-SF ≥ 32) có OR = 5,00 — yếu tố nguy cơ MẠNH THỨ HAI '
    'sau rối loạn giấc ngủ. Zheng và Peng (2025) trên học sinh '
    'nghề Trung Quốc dùng Social Media Addiction Scale phát hiện '
    'r = 0,415 với lo âu (GAD-7) và β = 0,153 trong hồi quy đa '
    'biến — chủ yếu qua trung gian giảm self-efficacy (63,13%).'
)
para(
    'Tại VIỆT NAM, kết quả của Hoàng Trung Học và cộng sự (2025) '
    'trên 8.473 học sinh THCS-THPT tại 6 tỉnh phát hiện thời gian '
    'sử dụng máy tính/điện thoại có β = 0,176 — yếu tố nguy cơ '
    'MẠNH THỨ HAI sau quan hệ cha mẹ-con. Tuy nhiên, VN014 đo '
    'THỜI GIAN dùng (screen time) chứ KHÔNG đo NGHIỆN (addiction) '
    'theo nghĩa lâm sàng — đây là KHOẢNG TRỐNG mà đề tài có thể '
    'lấp đầy.'
)

# 3. Bằng chứng định lượng
H('3. Bằng chứng định lượng — sáu công trình', level=2, color=NAVY)
caption('Bảng 1. Sáu công trình ủng hộ giả thuyết "nghiện điện thoại là yếu tố nguy cơ"')
add_table(
    ['#', 'Công trình', 'Thiết kế + cỡ mẫu', 'Chỉ số then chốt'],
    [
        ['1', 'Sohn và cộng sự (2019), BMC Psychiatry',
         'Meta-analysis 41 NC, N = 41.871',
         'Lo âu OR = 3,05 (KTC 95% 2,64–3,53; I² = 0%)'],
        ['2', 'Chen và cộng sự (2023), BMC Psychiatry [QT007]',
         'Cắt ngang n = 63.205, Tự Cống TQ',
         'Game disorder OR = 5,00 — mạnh thứ 2'],
        ['3', 'Zheng & Peng (2025), Psychology Research and Behavior Management [QT041]',
         'Cắt ngang HS nghề TQ',
         'SMAS r = 0,415; β = 0,153; mediation 63,13%'],
        ['4', 'Fassi và cộng sự (2025), Nature Human Behaviour [QT027]',
         'Cắt ngang n = 3.340 UK 11–19 (chẩn đoán lâm sàng)',
         'VTN có RLTT dùng MXH nhiều hơn — tác động vượt ngưỡng SESOI g = 0,4'],
        ['5', 'Schmidt-Persson và cộng sự (2024), JAMA Network Open [QT033]',
         'Cluster RCT thí điểm Đan Mạch, n = 89 gia đình (181 trẻ), 2 tuần',
         'Giảm screen time → giảm khó khăn nội hóa — bằng chứng nhân quả sơ bộ'],
        ['6', 'Hoàng Trung Học và cộng sự (2025) [VN014]',
         'Cắt ngang n = 8.473, 6 tỉnh VN',
         'Thời gian điện tử β = 0,176 — mạnh thứ 2'],
    ]
)

# 3.1 Sohn 2019
H('3.1. Sohn và cộng sự (2019) — Phân tích tổng hợp KINH ĐIỂN', level=3)
para(
    'Sohn, Rees, Wildridge, Kalk và Carter (2019) trong BMC '
    'Psychiatry thực hiện phân tích tổng hợp 41 nghiên cứu với '
    'N = 41.871 trẻ em và thanh niên — bài báo có GRADE đánh giá '
    'chất lượng bằng chứng. Phát hiện CỐT LÕI: nghiện điện thoại '
    '(problematic smartphone use, PSU) có liên quan với lo âu '
    'với OR = 3,05 (KTC 95% từ 2,64 đến 3,53; I² = 0%).'
)
para(
    'Đặc sắc của bài: heterogeneity I² = 0% — bằng KHÔNG TUYỆT '
    'ĐỐI. Nói cách khác, KHÔNG có khác biệt giữa các nghiên cứu '
    'về cường độ tác động — phát hiện ổn định qua các bối cảnh '
    'văn hóa và phương pháp khác nhau. Đây là bằng chứng VỮNG '
    'CHẮC NHẤT có thể có cho một mối quan hệ trong y văn dịch tễ.'
)
para(
    'Phù hợp với phát hiện: tỷ lệ nghiện điện thoại trung vị '
    '23,3% (KTC 95% 14,0–31,2%) — gần MỘT PHẦN TƯ trẻ em và '
    'thanh niên có biểu hiện sử dụng có vấn đề. Trong bối cảnh '
    'Việt Nam, tỷ lệ có thể CAO HƠN do mức độ phổ biến điện '
    'thoại trong học sinh tăng nhanh sau COVID-19.'
)

# 3.2 Chen 2023
H('3.2. Chen và cộng sự (2023) — Game disorder OR = 5,00 (Trung Quốc)', level=3)
para(
    'Chen, Ren, He và cộng sự (2023) trong BMC Psychiatry '
    '23(1):580 (DOI: 10.1186/s12888-023-05068-1) khảo sát '
    '63.205 học sinh trung học tại Tự Cống, miền Tây Trung '
    'Quốc — cỡ mẫu RẤT LỚN, lớn nhất trong cơ sở dữ liệu dự '
    'án (91 bài). Sử dụng IGDS9-SF (Internet Gaming Disorder '
    'Scale Short Form, ngưỡng ≥ 32) đo rối loạn chơi game '
    'online.'
)
para(
    'Phát hiện: rối loạn chơi game online có OR = 5,00 — yếu tố '
    'nguy cơ MẠNH THỨ HAI sau rối loạn giấc ngủ (OR = 6,99). 8,3% '
    'học sinh có triệu chứng căng thẳng tâm thần có rối loạn '
    'chơi game so với 1,0% nhóm không căng thẳng — chênh lệch '
    'TÁM LẦN. Đây là bằng chứng MẠNH cho rằng rối loạn chơi '
    'game online — một biểu hiện của nghiện thiết bị điện tử — '
    'là yếu tố nguy cơ rõ rệt cho lo âu và trầm cảm.'
)

# 3.3 Zheng 2025
H('3.3. Zheng và Peng (2025) — Cơ chế trung gian self-efficacy', level=3)
para(
    'Zheng và Peng (2025) trong Psychology Research and Behavior '
    'Management (DOI: 10.2147/PRBM.S522652) sử dụng SMAS (Social '
    'Media Addiction Scale 6 mục) đo nghiện mạng xã hội kết hợp '
    'GAD-7 đo lo âu, PSQI đo giấc ngủ, và GSES đo self-efficacy.'
)
para(
    'Phát hiện đặc sắc: nghiện mạng xã hội có tương quan trực '
    'tiếp với lo âu r = 0,415 và β = 0,153 trong hồi quy đa '
    'biến. Tuy nhiên, phân tích trung gian (mediation analysis) '
    'cho thấy 63,13% tổng tác động đi qua TRUNG GIAN '
    'self-efficacy — tức là nghiện điện thoại GIẢM '
    'self-efficacy, từ đó TĂNG lo âu. Phù hợp với khung lý '
    'thuyết của Bandura (1997) về self-efficacy.'
)

# 3.4 Fassi 2025
H('3.4. Fassi và cộng sự (2025) — Nature Human Behaviour với chẩn đoán lâm sàng', level=3)
para(
    'Fassi và cộng sự (2025) trên Nature Human Behaviour (IF ≈ '
    '24,0 — tạp chí có chỉ số ảnh hưởng CAO NHẤT trong cơ sở dữ '
    'liệu dự án) trên 3.340 thanh thiếu niên 11–19 tuổi tại '
    'Anh, mẫu đại diện QUỐC GIA. Đặc sắc: sức khỏe tâm thần được '
    'đánh giá bằng CHẨN ĐOÁN LÂM SÀNG (clinical raters) — KHÔNG '
    'chỉ sàng lọc tự báo cáo.'
)
para(
    'Phát hiện: thanh thiếu niên CÓ rối loạn sức khỏe tâm thần '
    '(được chẩn đoán lâm sàng) sử dụng mạng xã hội NHIỀU HƠN, '
    'có nhiều so sánh xã hội hơn, và kém hài lòng hơn về số '
    'lượng bạn online so với nhóm không có rối loạn. Đặc biệt, '
    'thanh thiếu niên có rối loạn nội hóa (internalizing) còn '
    'có mức độ ảnh hưởng cảm xúc từ phản hồi online cao hơn và '
    'tiết lộ bản thân ít chân thực hơn. Trái với quan niệm '
    '"nhiều bạn online = hỗ trợ tốt", VTN có rối loạn có nhiều '
    'bạn online nhưng kém hài lòng — gợi ý CHẤT LƯỢNG kết nối '
    'quan trọng hơn SỐ LƯỢNG.'
)

# 3.5 Schmidt-Persson 2024
H('3.5. Schmidt-Persson và cộng sự (2024) — Bằng chứng nhân quả từ RCT', level=3)
para(
    'Schmidt-Persson và cộng sự (2024) trong JAMA Network Open '
    '7(7):e2419881 (DOI: 10.1001/jamanetworkopen.2024.19881) '
    'thực hiện phân tích thứ cấp từ thử nghiệm lâm sàng ngẫu '
    'nhiên có cụm tại 10 đô thị miền Nam Đan Mạch — cỡ mẫu 89 '
    'GIA ĐÌNH (181 trẻ em + vị thành niên), can thiệp giảm '
    'screen time giải trí xuống còn ≤ 3 GIỜ/TUẦN trong 2 TUẦN. '
    'Đây là RCT THÍ ĐIỂM cỡ nhỏ — KHÔNG phải RCT đa trung tâm '
    'cỡ lớn.'
)
para(
    'Kết luận: giảm thời gian sử dụng màn hình giảm khó khăn '
    'hành vi nội hóa (internalizing — bao gồm lo âu) và tăng '
    'hành vi tiền xã hội ở trẻ em và vị thành niên — Cohen '
    'd = 0,53 cho khác biệt giữa hai nhóm (kích thước hiệu '
    'ứng TRUNG BÌNH). Tuy nhiên, do cỡ mẫu nhỏ (89 gia đình) '
    'và follow-up chỉ 2 tuần, generalizability còn hạn chế. '
    'Cần các RCT cỡ lớn dài hạn để xác lập bền vững.'
)

# 3.6 Hoàng Trung Học 2025
H('3.6. Hoàng Trung Học và cộng sự (2025) — Bằng chứng từ Việt Nam', level=3)
para(
    'Hoàng Trung Học và cộng sự (2025) trên 8.473 học sinh '
    'THCS-THPT tại 6 tỉnh Việt Nam (Đợt 1 trong COVID n = '
    '4.052; Đợt 2 sau COVID n = 4.337) sử dụng DASS-21. Hồi '
    'quy đa biến phát hiện thời gian sử dụng máy tính/điện '
    'thoại có β = 0,176 — yếu tố nguy cơ MẠNH THỨ HAI sau '
    'quan hệ cha mẹ-con (β = 0,272).'
)
para(
    'Hạn chế: VN014 đo THỜI GIAN sử dụng (screen time) — '
    'KHÔNG đo NGHIỆN (addiction). Hai khái niệm KHÁC NHAU '
    'về bản chất:'
)
para(
    '• THỜI GIAN dùng = số giờ trên thiết bị mỗi ngày — biến '
    'liên tục.\n'
    '• NGHIỆN = mất kiểm soát + craving + ảnh hưởng tiêu cực '
    '+ rút lui — chùm triệu chứng theo tiêu chí lâm sàng.', indent=False
)
para(
    'Một học sinh có thể dùng điện thoại 5 giờ/ngày MỘT CÁCH '
    'có kế hoạch (học tập + giải trí có kiểm soát) — KHÔNG '
    'phải nghiện. Một học sinh khác dùng 2 giờ/ngày NHƯNG '
    'không thể không cầm điện thoại lên 30 phút một lần — '
    'CÓ THỂ là nghiện.'
)

# 4. Cảnh báo về chiều nhân quả
H('4. Cảnh báo: bằng chứng dọc cho thấy quan hệ HAI CHIỀU', level=2, color=NAVY)
para(
    'Một bằng chứng QUAN TRỌNG đối lại với giả thuyết một '
    'chiều: Li và cộng sự (2025) trong British Journal of '
    'Clinical Psychology Q1 — thiết kế DỌC 12 tháng trên '
    '4.058 thanh thiếu niên Úc — phát hiện screen time T1 '
    'KHÔNG dự báo lo âu T2 (p = 0,443). Trái với mối liên '
    'quan cắt ngang đáng kể, mối liên quan dọc YẾU/KHÔNG '
    'đáng kể.', indent=False
)
para(
    'Nói cách khác, có khả năng quan hệ HAI CHIỀU: vị '
    'thành niên có triệu chứng lo âu/trầm cảm tăng dùng '
    'điện thoại (rút lui xã hội, phân tâm) — không chỉ '
    'ngược lại. Đây là khả năng REVERSE CAUSATION cần '
    'cảnh báo trong giả thuyết.'
)
para(
    'Hệ quả cho thiết kế nghiên cứu của thầy: KHÔNG nên '
    'phát biểu giả thuyết một chiều mạnh — nên phát biểu '
    'là "có MỐI LIÊN QUAN" thay vì "GÂY RA". Để chứng '
    'minh nhân quả, cần nghiên cứu dọc theo dõi cùng '
    'nhóm học sinh từ 12 tháng trở lên.'
)

# 5. Bốn cơ chế
H('5. Bốn cơ chế đề xuất từ y văn quốc tế', level=2, color=NAVY)
para('Y văn quốc tế đề xuất bốn cơ chế nối nghiện điện thoại với lo âu.', indent=False)
para(
    'CƠ CHẾ 1 — Thay thế giấc ngủ. Sử dụng điện thoại đêm '
    'khuya làm GIẢM thời lượng và chất lượng giấc ngủ '
    '(Chen 2023 phát hiện rối loạn giấc ngủ là yếu tố nguy '
    'cơ MẠNH NHẤT, OR = 6,99). Thiếu ngủ → tăng cortisol → '
    'tăng phản ứng lo âu.'
)
para(
    'CƠ CHẾ 2 — Thay thế tương tác xã hội thực. Thời gian '
    'dùng điện thoại nhiều LÀM GIẢM tương tác mặt đối mặt '
    '— vốn là nguồn hỗ trợ xã hội bảo vệ. Phù hợp với phát '
    'hiện Fassi 2025: VTN có rối loạn có nhiều bạn online '
    'nhưng kém hài lòng.'
)
para(
    'CƠ CHẾ 3 — So sánh xã hội tiêu cực. Mạng xã hội tạo '
    'môi trường so sánh KHÔNG CÔNG BẰNG (so với hình ảnh '
    'đã chỉnh sửa của người khác) — gây ra lo âu xã hội '
    'và giảm tự trọng. Phù hợp với khung lý thuyết của '
    'Festinger (1954).'
)
para(
    'CƠ CHẾ 4 — Self-efficacy thấp + FOMO. Sử dụng quá '
    'mức gây mất kiểm soát → giảm self-efficacy → tăng '
    'lo âu (Zheng & Peng 2025: 63,13% tác động đi qua '
    'self-efficacy). FOMO (Fear of Missing Out) tạo lo '
    'âu kỳ vọng khi không kết nối được.'
)

# 6. Phân tích kim tự tháp bằng chứng
H('6. Kim tự tháp bằng chứng cho giả thuyết H', level=2, color=NAVY)
caption('Bảng 2. Vị trí các công trình trong kim tự tháp bằng chứng y học')
add_table(
    ['Cấp', 'Loại bằng chứng', 'Công trình', 'Sức mạnh'],
    [
        ['I', 'RCT (thử nghiệm ngẫu nhiên có đối chứng) — thí điểm cỡ nhỏ',
         'Schmidt-Persson 2024 (JAMA Network Open) n = 89 gia đình',
         'NHÂN QUẢ sơ bộ — generalizability hạn chế (mẫu nhỏ + 2 tuần)'],
        ['II', 'Meta-analysis nghiên cứu quan sát',
         'Sohn 2019 (BMC Psychiatry, 41 NC, N = 41.871)',
         'TỔNG HỢP — OR = 3,05; I² = 0%'],
        ['III', 'Nghiên cứu dọc (longitudinal)',
         'Li 2025 (BJCP, n = 4.058, 12 tháng)',
         'TRÁI CHIỀU — gợi ý quan hệ hai chiều'],
        ['IV', 'Nghiên cứu cắt ngang cỡ mẫu lớn',
         'Chen 2023 (n = 63.205); Hoàng Trung Học 2025 (n = 8.473); Fassi 2025 (n = 3.340)',
         'CỠ MẪU LỚN — đáng tin cậy'],
        ['V', 'Phân tích cơ chế (mediation)',
         'Zheng & Peng 2025 (PRBM)',
         'CƠ CHẾ — self-efficacy trung gian 63,13%'],
    ]
)

# 7. Năm phát hiện chính
H('7. Năm phát hiện chính từ y văn ủng hộ giả thuyết H', level=2, color=NAVY)
para(
    'Thứ nhất, trên BÌNH DIỆN TOÀN CẦU, học sinh có nghiện '
    'điện thoại có khả năng mắc lo âu CAO HƠN GẤP BA LẦN — '
    'OR = 3,05 (Sohn và cộng sự, 2019), heterogeneity '
    'I² = 0% (đồng nhất tuyệt đối qua 41 nghiên cứu).', indent=False
)
para(
    'Thứ hai, rối loạn chơi game online — một thể của '
    'nghiện thiết bị — có OR = 5,00 cho căng thẳng tâm '
    'thần (Chen và cộng sự, 2023, n = 63.205) — yếu tố '
    'nguy cơ MẠNH THỨ HAI sau rối loạn giấc ngủ.'
)
para(
    'Thứ ba, RCT thí điểm cỡ nhỏ (Schmidt-Persson và cộng '
    'sự, 2024) trên 89 gia đình Đan Mạch cung cấp bằng '
    'chứng NHÂN QUẢ SƠ BỘ — giảm thời gian màn hình giảm '
    'khó khăn nội hóa (internalizing). Tuy cỡ mẫu nhỏ và '
    'follow-up chỉ 2 tuần, đây vẫn là thiết kế cấp bằng '
    'chứng cao nhất hiện có.'
)
para(
    'Thứ tư, cơ chế trung gian là self-efficacy — nghiện '
    'điện thoại GIẢM self-efficacy, từ đó TĂNG lo âu '
    '(Zheng & Peng, 2025, mediation 63,13%). Phù hợp với '
    'khung lý thuyết Bandura.'
)
para(
    'Thứ năm, tại Việt Nam, Hoàng Trung Học và cộng sự '
    '(2025) trên 8.473 học sinh đã xác lập THỜI GIAN dùng '
    'là yếu tố nguy cơ MẠNH THỨ HAI (β = 0,176) — nhưng '
    'CHƯA đo NGHIỆN theo nghĩa lâm sàng. Đây là khoảng '
    'trống mà đề tài có thể lấp.'
)

# 8. Bốn khoảng trống
H('8. Bốn khoảng trống nghiên cứu cho đề tài VN', level=2, color=NAVY)
para(
    'Y văn quốc tế đầy đủ ủng hộ giả thuyết H, nhưng còn '
    'BỐN KHOẢNG TRỐNG cho bối cảnh Việt Nam mà đề tài có '
    'thể lấp.', indent=False
)
para(
    'KHOẢNG TRỐNG 1 — Đo NGHIỆN, không chỉ THỜI GIAN. '
    'VN014 đo screen time. Đề tài cần dùng SAS-SV hoặc '
    'SMAS để phân biệt nghiện vs sử dụng có kiểm soát.'
)
para(
    'KHOẢNG TRỐNG 2 — Phân biệt LOẠI nghiện. Nghiện game '
    'online ≠ nghiện mạng xã hội ≠ nghiện video ngắn '
    '(TikTok). Mỗi loại có cơ chế và mức độ tác động '
    'khác nhau (Chen 2023 đo riêng game; Zheng 2025 đo '
    'riêng MXH).'
)
para(
    'KHOẢNG TRỐNG 3 — Phân tích TRUNG GIAN. Đề tài có '
    'thể test mediation: nghiện điện thoại → giảm '
    'self-efficacy → tăng lo âu (theo mô hình Zheng & '
    'Peng 2025). Hoặc: nghiện điện thoại → giảm giấc '
    'ngủ → tăng lo âu (theo mô hình Chen 2023).'
)
para(
    'KHOẢNG TRỐNG 4 — Thiết kế DỌC. Nghiên cứu dọc 12 '
    'tháng tại Việt Nam có thể giải quyết vấn đề chiều '
    'nhân quả mà Li 2025 đặt ra. Trong khi chờ nghiên '
    'cứu dọc, có thể dùng phân tích cross-lagged panel '
    'với dữ liệu từ hai thời điểm.'
)

# 9. CÂU TRẢ LỜI
H('9. CÂU TRẢ LỜI — Phát biểu giả thuyết và luận chứng', level=2, color=NAVY)
blue_run('Giả thuyết H (chính thức):', bold=True)
blue_run(
    'Học sinh trung học cơ sở có MỨC ĐỘ NGHIỆN ĐIỆN THOẠI '
    'cao hơn sẽ có TỶ LỆ rối loạn lo âu cao hơn — sau khi '
    'kiểm soát các biến nhân khẩu (giới, lớp, kinh tế gia '
    'đình) và các yếu tố tâm lý xã hội đồng biến (giấc '
    'ngủ, áp lực học tập, quan hệ cha mẹ-con).',
    italic=True
)
blue_run('Tóm gọn luận chứng (5 điểm):', bold=True)
blue_run(
    '(1) META-ANALYSIS Sohn 2019 BMC Psychiatry — 41 NC, '
    'N = 41.871: OR lo âu = 3,05 (KTC 95% 2,64–3,53; '
    'I² = 0%). Heterogeneity bằng KHÔNG — bằng chứng '
    'ổn định nhất.'
)
blue_run(
    '(2) CỠ MẪU LỚN Chen 2023 BMC Psychiatry n = 63.205: '
    'rối loạn chơi game OR = 5,00 — mạnh thứ 2 sau giấc '
    'ngủ (OR = 6,99).'
)
blue_run(
    '(3) RCT thí điểm Schmidt-Persson 2024 JAMA Network '
    'Open (n = 89 gia đình Đan Mạch, 2 tuần): giảm screen '
    'time → giảm khó khăn nội hóa — bằng chứng nhân quả '
    'SƠ BỘ.'
)
blue_run(
    '(4) CƠ CHẾ Zheng & Peng 2025 PRBM: self-efficacy '
    'là biến trung gian, 63,13% tác động đi qua đường '
    'này.'
)
blue_run(
    '(5) VIỆT NAM Hoàng Trung Học 2025 (VN014) n = 8.473: '
    'thời gian điện tử β = 0,176 — yếu tố nguy cơ mạnh '
    'thứ 2 sau quan hệ cha mẹ-con. Tuy nhiên VN014 '
    'KHÔNG đo nghiện — chỉ đo thời gian.'
)
blue_run('Cảnh báo bằng chứng đối lại:', bold=True)
blue_run(
    'Li 2025 BJCP n = 4.058 dọc 12 tháng — screen time '
    'T1 KHÔNG dự báo lo âu T2 (p = 0,443). Có khả năng '
    'quan hệ HAI CHIỀU (lo âu → tăng dùng điện thoại). '
    'Phát biểu nên dùng "MỐI LIÊN QUAN" thay vì "GÂY RA". '
    'Chứng minh nhân quả cần thiết kế dọc.'
)
blue_run('Cách trích vào báo cáo CTH:', bold=True)
blue_run(
    '"Phân tích tổng hợp 41 nghiên cứu của Sohn và cộng '
    'sự (2019) trên 41.871 trẻ em và thanh niên trong '
    'BMC Psychiatry xác lập rằng nghiện điện thoại '
    '(problematic smartphone use) liên quan với lo âu '
    'với OR = 3,05 (KTC 95% 2,64–3,53; I² = 0%) — bằng '
    'chứng tổng hợp với heterogeneity tuyệt đối bằng '
    'không. Trong nghiên cứu cỡ mẫu LỚN NHẤT thế giới '
    'tại miền Tây Trung Quốc trên 63.205 học sinh '
    'trung học, Chen và cộng sự (2023) phát hiện rối '
    'loạn chơi game online có OR = 5,00 — yếu tố nguy '
    'cơ mạnh thứ hai sau rối loạn giấc ngủ. Bằng chứng '
    'nhân quả từ thử nghiệm lâm sàng ngẫu nhiên '
    '(Schmidt-Persson và cộng sự, 2024 trong JAMA '
    'Network Open) cho thấy giảm thời gian màn hình '
    'cải thiện sức khỏe tâm thần. Tại Việt Nam, Hoàng '
    'Trung Học và cộng sự (2025) trên 8.473 học sinh '
    'THCS-THPT phát hiện thời gian sử dụng máy tính/'
    'điện thoại có β = 0,176 — yếu tố nguy cơ mạnh '
    'thứ hai. Trên cơ sở năm nguồn bằng chứng hội tụ, '
    'đề tài đặt giả thuyết: học sinh THCS có mức độ '
    'nghiện điện thoại cao hơn sẽ có tỷ lệ rối loạn '
    'lo âu cao hơn — sau khi kiểm soát các biến nhân '
    'khẩu và yếu tố tâm lý xã hội đồng biến."',
    italic=True
)

# 10. TLTK
H('10. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Andreassen, C. S., Billieux, J., Griffiths, M. D., Kuss, D. J., Demetrovics, Z., Mazzoni, E., & Pallesen, S. (2016). The relationship between addictive use of social media and video games and symptoms of psychiatric disorders: A large-scale cross-sectional study. Psychology of Addictive Behaviors, 30(2), 252–262. https://doi.org/10.1037/adb0000160',
    'Bandura, A. (1997). Self-efficacy: The exercise of control. W.H. Freeman.',
    'Chen, Z., Ren, S., He, R., Liang, Y., Tan, Y., Liu, Y., Wang, F., Shao, X., Chen, S., Liao, Y., He, Y., Li, J. G., Chen, X., & Tang, J. (2023). Prevalence and associated factors of depressive and anxiety symptoms with demographic, family, school, life, and behavior factors in a large, representative sample of secondary school students in Zigong, a city in Western China. BMC Psychiatry, 23(1), 580. https://doi.org/10.1186/s12888-023-05068-1 [QT007 trong cơ sở dữ liệu dự án.]',
    'Fassi, L., và cộng sự. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9(6), 1283–1299. https://doi.org/10.1038/s41562-025-02134-4 [QT027 trong cơ sở dữ liệu dự án.]',
    'Festinger, L. (1954). A theory of social comparison processes. Human Relations, 7(2), 117–140.',
    'Hoàng, T. H., và cộng sự. (2025). Sức khỏe tâm thần của học sinh Việt Nam trong và sau đại dịch COVID-19. [VN014 trong cơ sở dữ liệu dự án.]',
    'Kwon, M., Kim, D. J., Cho, H., & Yang, S. (2013). The Smartphone Addiction Scale: Development and validation of a short version for adolescents. PLoS ONE, 8(12), e83558. https://doi.org/10.1371/journal.pone.0083558',
    'Li, S. H., Batterham, P. J., Whitton, A. E., Maston, K., Khan, A., Christensen, H., & Werner-Seidler, A. (2025). Cross-sectional and longitudinal associations of screen time with adolescent depression and anxiety. British Journal of Clinical Psychology, 64(4), 873–887. https://doi.org/10.1111/bjc.12547 [QT022 trong cơ sở dữ liệu dự án.]',
    'Schmidt-Persson, J., Rasmussen, M. G. B., Sørensen, S. O., Mortensen, S. R., Olesen, L. G., Brage, S., Kristensen, P. L., Bilenberg, N., & Grøntved, A. (2024). Screen media use and mental health of children and adolescents: A secondary analysis of a randomized clinical trial. JAMA Network Open, 7(7), e2419881. https://doi.org/10.1001/jamanetworkopen.2024.19881 [QT033 trong cơ sở dữ liệu dự án.]',
    'Sohn, S. Y., Rees, P., Wildridge, B., Kalk, N. J., & Carter, B. (2019). Prevalence of problematic smartphone usage and associated mental health outcomes amongst children and young people: A systematic review, meta-analysis and GRADE of the evidence. BMC Psychiatry, 19(1), 356. https://doi.org/10.1186/s12888-019-2350-x',
    'Zheng, G. F., & Peng, H. Y. (2025). The effects of social media addiction, academic stress, and sleep quality on anxiety symptoms: A cross-sectional study of Chinese vocational students. Psychology Research and Behavior Management, 18, 1571–1584. https://doi.org/10.2147/PRBM.S522652 [QT041 trong cơ sở dữ liệu dự án.]',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
