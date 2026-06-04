# -*- coding: utf-8 -*-
"""Sua not loi quan trong cuoi cung: 1.2.4.3 trung (para 529 + 544).
Va kiem tra format con sot.
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')


def fix_runs(p, replacements):
    changed = False
    for run in p.runs:
        for old, new in replacements:
            if old in run.text:
                run.text = run.text.replace(old, new)
                changed = True
    return changed


d = Document(FILE)
stats = {'fixed': 0, 'cascaded': 0}

# =========================================================
# Para 529 + 544: 1.2.4.3 trung lap
# Cau truc logic:
# - 1.2.4.1 (508): HS THCS (dinh nghia)
# - 1.2.4.2 (524): Suc khoe tam ly HS THCS
# - 1.2.4.3 (529): HS THCS va RLLA (intro/quan he) - GIU
# - 1.2.4.3 (544) -> 1.2.4.4: RLLA o HS THCS (mo ta chung)
# - 1.2.4.4 (549) -> 1.2.4.5: Bieu hien RLLA o HS THCS
# =========================================================
print("[1] Sua duplicate 1.2.4.3 (para 544) + cascade...")

# Para 544: "1.2.4.3." -> "1.2.4.4."
if 544 < len(d.paragraphs):
    p = d.paragraphs[544]
    if '1.2.4.3.' in p.text:
        if fix_runs(p, [('1.2.4.3.', '1.2.4.4.')]):
            stats['fixed'] += 1
            print(f"  Para 544: '1.2.4.3.' -> '1.2.4.4.' (Rối loạn lo âu ở HS THCS)")

# Para 549: "1.2.4.4." -> "1.2.4.5."
if 549 < len(d.paragraphs):
    p = d.paragraphs[549]
    if '1.2.4.4.' in p.text:
        if fix_runs(p, [('1.2.4.4.', '1.2.4.5.')]):
            stats['cascaded'] += 1
            print(f"  Para 549: '1.2.4.4.' -> '1.2.4.5.' (Biểu hiện RLLA - cascade)")


# =========================================================
# Sua not so 6.  7.  8.  10.  (double space sau so)
# =========================================================
print("\n[2] Sua double space trong heading number...")
heading_paras = [172, 176, 183, 194, 203]
for idx in heading_paras:
    if idx >= len(d.paragraphs): continue
    p = d.paragraphs[idx]
    # Replace patterns like "6.  N" -> "6. N"
    for run in p.runs:
        new_text = re.sub(r'(\d+\.)  +', r'\1 ', run.text)
        if new_text != run.text:
            run.text = new_text
            stats['fixed'] += 1
            print(f"  Para {idx}: fix double space after number")


# =========================================================
# Verify 5 lỗi ưu tiên đã sửa
# =========================================================
print("\n[3] Verify 5 loi uu tien cao...")

checks = [
    (149, 'cs.', 'cs..', 'Para 149 (double dot fixed)'),
    (331, '1.1.6.5', '1.2.6.5', 'Para 331 (1.2.6.5 -> 1.1.6.5)'),
    (544, '1.2.4.4', '1.2.4.3', 'Para 544 (duplicate fixed)'),
    (971, '3.4.2.1', None, 'Para 971 (kept as original)'),
    (1002, '3.4.2.4', '3.4.2.1', 'Para 1002 (3.4.2.1 -> 3.4.2.4)'),
    (1024, '3.4.3.2', None, 'Para 1024 (kept as original)'),
    (1036, '3.4.3.3', '3.4.3.2', 'Para 1036 (3.4.3.2 -> 3.4.3.3)'),
]
for idx, want, not_want, desc in checks:
    if idx >= len(d.paragraphs): continue
    txt = d.paragraphs[idx].text[:80]
    has_want = want in txt
    has_not_want = (not_want is None) or (not_want not in txt or txt.count(want) > 0)
    mark = '✓' if (has_want and (not_want is None or not_want not in txt[:len(not_want)+5])) else '?'
    print(f"  {mark} {desc}: {txt!r}")

d.save(FILE)
print(f"\nSaved: {FILE}")
print(f"\nFixed: {stats['fixed']}, Cascaded: {stats['cascaded']}")
