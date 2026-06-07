"""Final round: verify N=1,352 in both, confirm β consistency, anti-fab refs."""
import re
import json
from pathlib import Path
from docx import Document

ROOT = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au")
Q1V7 = ROOT / "bai-bao-Q1" / "Draft_Q1_SongNgu_v7_01062026.docx"
Q2V1 = ROOT / "bai-bao-Q1" / "Draft_Q2_SongNgu_v1_07062026.docx"


def text_of(path):
    doc = Document(str(path))
    parts = []
    for p in doc.paragraphs:
        if p.text.strip():
            parts.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        parts.append(p.text)
    return "\n".join(parts)


def main():
    out = {}
    for label, path in [("Q1v7", Q1V7), ("Q2v1", Q2V1)]:
        text = text_of(path)
        out[label] = {
            "N_1352_EN_count": len(re.findall(r"1,352", text)),
            "N_1352_VN_count": len(re.findall(r"1\.352", text)),
            "Male_614_count": len(re.findall(r"\b614\b", text)),
            "Female_738_count": len(re.findall(r"\b738\b", text)),
            "M_45_4_count": len(re.findall(r"45[\.,]4", text)),
            "F_54_6_count": len(re.findall(r"54[\.,]6", text)),
            "beta_0_376": len(re.findall(r"0[\.,]376", text)),
            "beta_0_510": len(re.findall(r"0[\.,]510", text)),
            "beta_0_669": len(re.findall(r"0[\.,]669", text)),
            "beta_0_215": len(re.findall(r"0[\.,]215", text)),
            "beta_0_253": len(re.findall(r"0[\.,]253", text)),
            "GAD_F_44_48": "44.48" in text or "44,48" in text,
            "SocAD_F_45_98": "45.98" in text or "45,98" in text,
            "SAD_F_0_246": "0.246" in text or "0,246" in text,
            "GAD_59_47": "59.47" in text or "59,47" in text,
            "Male_GAD_51_43": "51.43" in text or "51,43" in text,
            "SocAD_52_74": "52.74" in text or "52,74" in text,
            "Male_SocAD_43_20": "43.20" in text or "43,20" in text,
            "SAD_F_24_76": "24.76" in text or "24,76" in text,
            "Male_SAD_25_42": "25.42" in text or "25,42" in text,
            "Cronbach_70": "0.70" in text and "0,70" in text,
            "CFI_90": "0.90" in text and "0,90" in text,
            "RMSEA_08": "0.08" in text and "0,08" in text,
            "ASEAN_prev_11_9": "11.9" in text and "11,9" in text,
            "VN_prev_10_1": "10.1" in text and "10,1" in text,
            "Grade_368": text.count("368"),
            "Grade_316": text.count("316"),
            "Grade_340": text.count("340"),
            "Grade_328": text.count("328"),
        }
    # Print summary table
    keys = list(out["Q1v7"].keys())
    print(f"{'KEY':<28} {'Q1v7':>12} {'Q2v1':>12}")
    for k in keys:
        print(f"{k:<28} {str(out['Q1v7'][k]):>12} {str(out['Q2v1'][k]):>12}")

    json.dump(out, open(ROOT/"06_Scripts"/"_verify_q2polish_finalize.json","w",encoding="utf-8"), ensure_ascii=False, indent=2)


if __name__ == "__main__":
    main()
