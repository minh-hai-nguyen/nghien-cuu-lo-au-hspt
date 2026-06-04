# -*- coding: utf-8 -*-
"""Generate v2: KTC vs CrI — revised with psychology methodology context."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Giai_thich_KTC_vs_CrI_cho_thay_v2.docx')

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False, underline=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if underline: r.underline = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def PRun(parts):
    p = doc.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 12))
        if opts.get('bold'): r.bold = True
        if opts.get('italic'): r.italic = True
        if opts.get('underline'): r.underline = True
        if opts.get('red'): r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif opts.get('color'): r.font.color.rgb = opts['color']
    return p

def H1(t): return P(t, bold=True, size=16, color=RGBColor(0x1F, 0x3A, 0x68), align=WD_ALIGN_PARAGRAPH.CENTER)
def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

def note_box(text, color_hex='FFF3E0'):
    """Warning/note box with colored background."""
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    c.text = ''
    r = c.paragraphs[0].add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10.5)
    r.italic = True
    set_cell_bg(c, color_hex)


# ============================================================
# TIÊU ĐỀ
# ============================================================
H1('GIẢI THÍCH KÝ HIỆU KTC VÀ CrI')
H1('TRONG BỐI CẢNH NGHIÊN CỨU TÂM LÝ HỌC')
P('Phiên bản 2 — đã rà soát phê phán theo chuẩn APA 7 + phương pháp luận MA tâm lý học',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P('Người soạn: Nhóm dự án Lo âu | Tháng 04/2026',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P()

note_box(
    'Phiên bản v2 (15/04/2026) đã sửa 5 điểm so với v1: (1) Phân biệt MD vs SMD rõ ràng — trong tâm lý '
    'học lâm sàng khi đo bằng thang khác nhau BẮT BUỘC dùng SMD; (2) Thực tế APA/tâm lý học vẫn diễn giải '
    'CI "lỏng lẻo" như CrI — đúng về implication, sai về philosophy; (3) CrI vs CI phụ thuộc chất lượng '
    'prior — với prior non-informative, hai kết quả gần như trùng; (4) SMD trong Cochrane MA là Hedges\' g '
    '(đã hiệu chỉnh small-sample), không phải Cohen\'s d gốc; (5) Bổ sung chuẩn APA 7, random-effects, '
    'và TOST/SESOI.',
    color_hex='FFFDE7')
P()

# ============================================================
# PHẦN 1
# ============================================================
H2('1. KTC (Khoảng tin cậy) — thuật ngữ tiếng Việt chung')

PRun([
    ('KTC', {'bold': True}),
    (' = ', {}),
    ('K', {'bold': True, 'underline': True}), ('hoảng ', {}),
    ('T', {'bold': True, 'underline': True}), ('in ', {}),
    ('C', {'bold': True, 'underline': True}), ('ậy. Tiếng Anh có ', {}),
    ('hai', {'bold': True}), (' khái niệm khác nhau về nền tảng thống kê:', {}),
])

P('• CI = Confidence Interval — Khoảng tin cậy theo trường phái tần suất (frequentist)', size=12)
P('• CrI = Credible Interval — Khoảng tin cậy theo trường phái Bayesian', size=12)
P()
PRun([
    ('Vấn đề dịch thuật: ', {'bold': True, 'italic': True}),
    ('tiếng Việt chỉ có ', {}),
    ('một ', {'bold': True}),
    ('cụm từ "khoảng tin cậy" nhưng tiếng Anh có ', {}),
    ('hai', {'bold': True}),
    (' khái niệm. Khi ghi "KTC 95 %" trong báo cáo, nên đi kèm "CI" hoặc "CrI" để tránh nhầm lẫn.', {}),
])
P()
note_box(
    'TRONG BỐI CẢNH TÂM LÝ HỌC: APA 7th edition (chuẩn viết bài tâm lý học quốc tế) khuyến nghị báo cáo '
    'ước lượng điểm kèm CI 95 % cho MỌI kết quả chính. Bayesian analysis đang tăng trong tâm lý học '
    '(đặc biệt Bayesian meta-analysis sau 2018) nhưng CI tần suất vẫn là chuẩn mặc định trong đa số '
    'tạp chí (JAMA, Lancet Psychiatry, J Clin Child Adolesc Psychol, Behav Res Ther...). Do đó khi đọc '
    'bài nếu KHÔNG THẤY ghi "Bayesian" thì mặc định là CI tần suất.',
    color_hex='E3F2FD')
P()

# ============================================================
# PHẦN 2
# ============================================================
H2('2. So sánh CI vs CrI (trong bối cảnh tâm lý học)')

MakeTable(
    ['Đặc điểm', 'CI (Confidence Interval)', 'CrI (Credible Interval)'],
    [
        ('Tên tiếng Việt', 'Khoảng tin cậy (tần suất)', 'Khoảng tin cậy Bayesian'),
        ('Trường phái', 'Frequentist (tần suất)', 'Bayesian'),
        ('Phổ biến trong tâm lý học',
         'Chuẩn mặc định — 85–90 % bài báo hiện tại dùng CI',
         'Tăng nhanh từ 2018, chủ yếu trong MA/NMA, Bayesian ANOVA'),
        ('Diễn giải ĐÚNG VỀ LÝ THUYẾT',
         '"Nếu lặp lại NC vô hạn lần, 95 % các khoảng tính được sẽ chứa giá trị thật" — một tuyên bố về PROCEDURE',
         '"Có 95 % xác suất giá trị thật nằm trong khoảng này" — tuyên bố về xác suất posterior'),
        ('Diễn giải THỰC TIỄN trong tâm lý học',
         'Nhiều nhà TLH (và cả bài báo lâm sàng) diễn giải LOOSE như CrI — về thực tiễn acceptable dù về triết học không chính xác',
         'Giống với diễn giải loose của CI — trong đa số báo cáo phương sai bình thường, hai kết quả gần như trùng số'),
        ('Giả định về mẫu',
         'Cần n đủ lớn để xấp xỉ phân phối chuẩn (đặc biệt cho t-test, CI quanh mean)',
         'Ít phụ thuộc mẫu lớn HƠN — nhưng với n nhỏ + prior non-informative vẫn cho posterior rộng/noisy'),
        ('Vai trò của prior',
         'Không có — mỗi NC độc lập',
         'Quan trọng: prior informative + data cho hậu nghiệm. Prior non-informative (chuẩn default) gần như trùng với CI'),
        ('Khi nào KHÁC BIỆT rõ?',
         '—',
         'Khi (a) prior informative mạnh; (b) n rất nhỏ (< 30); (c) phân phối không chuẩn. Trong đa số tâm lý học MA (n tổng lớn, prior weakly informative) → CI và CrI trùng số gần như tuyệt đối'),
    ])
P()

# ============================================================
# PHẦN 3 — DIỄN GIẢI VÍ DỤ
# ============================================================
H2('3. Ví dụ minh hoạ: SMD (KTC 95 % CrI) = –0,19 [–0,22; –0,17]')

note_box(
    'LƯU Ý SỬA TỪ V1: Ví dụ này dùng SMD (Standardized Mean Difference — khác biệt trung bình chuẩn hoá) '
    'chứ KHÔNG phải MD (Mean Difference). Trong meta-analysis tâm lý học, khi các NC dùng các thang '
    'đo khác nhau (GAD-7, SCARED, RCADS, STAI-C...) BẮT BUỘC dùng SMD để chuẩn hoá. Sử dụng MD chỉ khi '
    'MỌI NC đều dùng CÙNG một thang đo (hiếm gặp trong MA lớn).',
    color_hex='FFEBEE')
P()

H3('Cách đọc TRIẾT HỌC ĐÚNG (Bayesian CrI):')
P('"Có 95 % xác suất rằng hiệu quả chuẩn hoá thực sự (SMD) nằm trong khoảng –0,22 đến –0,17. '
  'Ước lượng điểm tốt nhất là –0,19."',
  italic=True, color=RGBColor(0x0B, 0x5E, 0x20))
P()

H3('Cách đọc CỦA NHÀ LÂM SÀNG (thực tế):')
P('"Hiệu quả có độ lớn khoảng –0,19 chuẩn (nghĩa là nhóm can thiệp giảm 0,19 độ lệch chuẩn so với đối chứng), '
  'với khoảng ước lượng hợp lý từ –0,22 đến –0,17. Không trùng 0 → hiệu quả thực sự tồn tại."',
  italic=True, color=RGBColor(0x0B, 0x5E, 0x20))
P()

note_box(
    'TRONG TÂM LÝ HỌC THỰC HÀNH: Về cơ bản, hai cách đọc trên đều dẫn đến cùng một kết luận LÂM SÀNG '
    '(effect size nhỏ, có ý nghĩa thống kê nhưng không đủ ý nghĩa lâm sàng). Khác biệt triết học giữa '
    'CI vs CrI có ý nghĩa học thuật quan trọng nhưng KHÔNG thay đổi decision lâm sàng. Nhiều nhà TLH '
    'chuyển từ frequentist sang Bayesian chủ yếu vì: (1) Bayesian NMA xếp hạng can thiệp dễ hơn (SUCRA); '
    '(2) tích hợp bằng chứng đa nguồn; (3) không cần bác bỏ null, cho phép quantify evidence FOR null.',
    color_hex='E8F5E9')
P()

H3('Ý nghĩa lâm sàng — áp dụng chuẩn Cohen cho SMD:')
MakeTable(
    ['Độ lớn |SMD|', 'Mô tả', 'Hàm ý lâm sàng'],
    [
        ('< 0,2', 'Trivial (không đáng kể)', 'Không có ý nghĩa lâm sàng dù có ý nghĩa thống kê'),
        ('0,2–0,5', 'Small (nhỏ)', 'Hiệu quả nhỏ — cần cân nhắc cost-benefit'),
        ('0,5–0,8', 'Medium (trung bình)', 'Hiệu quả đáng kể — có thể triển khai'),
        ('> 0,8', 'Large (lớn)', 'Hiệu quả mạnh — ưu tiên triển khai'),
    ])
P()
PRun([
    ('Áp dụng cho Zhang 2026 (SMD = –0,19): ', {'bold': True}),
    ('|SMD| < 0,2 → TRIVIAL theo Cohen. Có ý nghĩa thống kê (CrI không chứa 0) nhưng ', {}),
    ('ý nghĩa lâm sàng rất hạn chế', {'bold': True}),
    ('. Đây là lý do tác giả khuyến cáo ', {}),
    ('KHÔNG universal triển khai', {'italic': True}),
    (' mà nên chọn targeted (cho HS triệu chứng).', {}),
])
P()

# ============================================================
# PHẦN 4 — KÝ HIỆU CHUẨN (chi tiết hơn)
# ============================================================
H2('4. Quy ước ký hiệu — chuẩn tâm lý học và Cochrane')

H3('4.1. Khác biệt MD vs SMD — CẨN THẬN khi đọc báo cáo')
MakeTable(
    ['Ký hiệu', 'Tên đầy đủ', 'Dùng khi nào', 'Ví dụ trong tâm lý học'],
    [
        ('MD', 'Mean Difference (khác biệt trung bình thô)',
         'TẤT CẢ các NC dùng CÙNG một thang đo (cùng scale, cùng range)',
         'MA các RCT dùng chung GAD-7 → MD = –2,3 điểm (trên thang 0–21)'),
        ('SMD', 'Standardized Mean Difference (khác biệt chuẩn hoá = d/g)',
         'Các NC dùng thang đo KHÁC NHAU (đo cùng construct nhưng scale khác)',
         'MA các RCT dùng đa dạng (GAD-7 + SCARED + RCADS) → SMD = –0,19 (không đơn vị)'),
    ])
P('Công thức: ', bold=True, italic=True)
P('• MD = mean(can thiệp) – mean(đối chứng)', size=10)
P('• SMD = MD ÷ SD_pooled (SMD không có đơn vị, có thể so sánh xuyên scales)', size=10)
P()

H3('4.2. Cohen\'s d vs Hedges\' g — CẢ HAI đều là SMD nhưng khác nhau')
MakeTable(
    ['Ký hiệu', 'Công thức', 'Khi nào dùng'],
    [
        ('Cohen\'s d', 'd = (M1 – M2) / SD_pooled',
         'Công thức gốc Cohen 1988. Lệch dương nhẹ khi n nhỏ (biased cao hơn thực tế)'),
        ('Hedges\' g', 'g = d × J(df) trong đó J(df) = 1 – 3/(4·df – 1)',
         'Hedges-Olkin 1985: hiệu chỉnh small-sample bias. CHUẨN COCHRANE cho meta-analysis'),
    ])
note_box(
    'QUAN TRỌNG: Trong báo cáo MA tâm lý học hiện đại (Cochrane, PRISMA), SMD mặc định là Hedges\' g '
    '(không phải Cohen\'s d gốc). Khi đọc bài, nếu ghi chỉ "SMD" mà không nói rõ → nhiều khả năng là '
    'Hedges\' g. Sự khác biệt giữa d và g thường < 1–2 % với n > 50, nhưng quan trọng về mặt chuẩn '
    'báo cáo. Tiêu chí Cohen (< 0,2 trivial...) áp dụng cho CẢ d và g.',
    color_hex='FFF3E0')
P()

H3('4.3. Các ký hiệu thống kê khác trong tâm lý học')
MakeTable(
    ['Ký hiệu', 'Ý nghĩa', 'Thang tiêu chí'],
    [
        ('OR [KTC ...]', 'Odds Ratio (tỷ số odds)',
         'OR = 1: không effect; OR > 1: tăng nguy cơ; OR < 1: bảo vệ. Trong psych: OR 2–3 = medium'),
        ('RR [KTC ...]', 'Risk Ratio (tỷ số nguy cơ tương đối)',
         'Giống OR khi tỷ lệ biến cố thấp; diễn giải trực quan hơn'),
        ('HR [KTC ...]', 'Hazard Ratio (tỷ số nguy cơ tức thời)',
         'Dùng trong survival analysis (relapse, drop-out theo thời gian)'),
        ('β [KTC ...]', 'Hệ số hồi quy chuẩn hoá',
         'Trong SEM/multiple regression: β = 0,1 small; 0,3 medium; 0,5 large (Cohen)'),
        ('r [KTC ...]', 'Hệ số tương quan Pearson',
         'Cohen: r = 0,1 small; 0,3 medium; 0,5 large'),
        ('η²p (eta squared partial)', 'Hiệu ứng trong ANOVA',
         'η² = 0,01 small; 0,06 medium; 0,14 large'),
        ('NNT', 'Number Needed to Treat',
         'Số HS cần điều trị để 1 HS có cải thiện. NNT 10–20 = acceptable cho prevention'),
        ('SUCRA',
         'Surface Under Cumulative Ranking curve (NMA Bayesian)',
         'SUCRA 0–1 (càng cao càng tốt). > 0,7 = xếp hạng cao'),
    ])
P()

# ============================================================
# PHẦN 5 — TRONG THƯ VIỆN
# ============================================================
H2('5. Khi nào CrI xuất hiện trong thư viện dự án?')

P('Bảng đầy đủ các bài Bayesian trong thư viện 68 papers của dự án Lo âu:', italic=True)
P()

MakeTable(
    ['Canonical', 'Tác giả / Năm', 'Thiết kế', 'Kết quả chính (dùng CrI)'],
    [
        ('QT049', 'Zhang, Liang & Kang 2026',
         'Bayesian hierarchical MA',
         'SMD anxiety –0,19 (95 % CrI: –0,22; –0,17); SMD depression –0,06 (–0,08; –0,04)'),
        ('QT029', 'Li et al. 2025 (BMC Psychiatry)',
         'Bayesian NMA 30 RCT',
         'CBT cho lo âu trẻ em: SUCRA 0,66 (xếp hạng cao); PE SUCRA 0,51'),
        ('QT039', 'Xian, Zhang & Jiang 2024 (J Affect Disord)',
         'Bayesian NMA 30 RCT cho SAD',
         'iCBT: SUCRA 71,2 % (hạng 1); gCBT: SUCRA 68,4 %'),
    ])
P()

note_box(
    'TRONG PSYCH MA THỰC TẾ: Cả 3 bài Bayesian trên đều dùng prior weakly informative '
    '(vague prior — chuẩn mặc định). Do đó kết quả CrI của họ gần như trùng với kết quả '
    'CI frequentist nếu áp dụng cùng data. Điểm ưu thế thực sự của Bayesian NMA (QT029, '
    'QT039) là cho phép xếp hạng SUCRA — một giá trị dễ truyền thông cho chính sách ("iCBT '
    'xếp thứ nhất với 71 % probability of being best").',
    color_hex='E1F5FE')
P()

# ============================================================
# PHẦN 6 — CHUẨN APA 7
# ============================================================
H2('6. Chuẩn APA 7 cho báo cáo thống kê trong tâm lý học')

P('APA Publication Manual (7th ed., 2020) — quy tắc bắt buộc:', italic=True)
P()
P('1. ', bold=True)
PRun([
    ('Báo cáo CẢ ước lượng điểm VÀ CI 95 %', {'bold': True}),
    (' cho mọi kết quả chính. Ví dụ:', {}),
])
P('   ✅ "M = 3,42, 95 % CI [2,85; 3,99]"', italic=True, color=RGBColor(0x0B, 0x5E, 0x20))
P('   ❌ "M = 3,42, p < 0,05" (không đủ thông tin — điểm không kèm CI)', italic=True, color=RGBColor(0xC0, 0x00, 0x00))

P('2. ', bold=True)
PRun([
    ('Effect size là BẮT BUỘC', {'bold': True}),
    (': báo cáo Cohen\'s d, η², r, OR hoặc tương đương.', {}),
])

P('3. ', bold=True)
PRun([
    ('Cho Bayesian analysis', {'bold': True}),
    (': cần ghi rõ prior, báo cáo 95 % CrI, posterior median hoặc mean, và BF (Bayes Factor) nếu có.', {}),
])

P('4. ', bold=True)
PRun([
    ('Random-effects là chuẩn mặc định', {'bold': True}),
    (' cho meta-analysis tâm lý học. Fixed-effects chỉ dùng khi heterogeneity thấp (I² < 25 %). Báo cáo τ², I² cùng CI.', {}),
])

P('5. ', bold=True)
PRun([
    ('CI/CrI nên cho điểm ước lượng của tất cả tests hậu nghiệm', {'bold': True}),
    (' (post-hoc comparisons), không chỉ tests chính.', {}),
])
P()

# ============================================================
# PHẦN 7 — TOST/SESOI
# ============================================================
H2('7. Equivalence Testing (TOST) và SESOI trong tâm lý học')

P('Một xu hướng mới (từ ~2017) trong tâm lý học: thay vì chỉ test "có khác 0 hay không", các nhà '
  'nghiên cứu dùng ', italic=True)
PRun([
    ('Equivalence Testing (TOST = Two One-Sided Tests)', {'bold': True, 'italic': True}),
    (' để kết luận "hiệu ứng nhỏ đến mức không có ý nghĩa lâm sàng". Điểm chốt: ', {'italic': True}),
])
P()

P('• ', bold=True)
PRun([
    ('SESOI', {'bold': True}),
    (' = Smallest Effect Size of Interest (cỡ hiệu ứng nhỏ nhất đáng quan tâm). Ví dụ trong tâm lý '
     'trẻ em lo âu: SESOI = ±0,2 (nhỏ nhất mà thực tế lâm sàng quan tâm).', {}),
])
P('• ', bold=True)
PRun([
    ('TOST', {'bold': True}),
    (': nếu 90 % CI nằm TOÀN BỘ trong khoảng [–SESOI, +SESOI] → KẾT LUẬN TƯƠNG ĐƯƠNG (không có hiệu ứng có ý nghĩa lâm sàng).', {}),
])
P()

note_box(
    'Áp dụng TOST cho Zhang 2026 (SMD = –0,19, 95 % CrI: –0,22 đến –0,17): với SESOI = 0,2, CrI có '
    'upper bound –0,17 (trong khoảng equivalence [–0,2; +0,2]) nhưng lower bound –0,22 (NGOÀI '
    'khoảng). → KHÔNG kết luận được tương đương. Nói cách khác: hiệu ứng vừa đủ ngoài trivial (có '
    'ý nghĩa thống kê) nhưng không đủ lớn để có ý nghĩa lâm sàng. Đây chính xác là vùng "xám" khó '
    'quyết định mà tác giả phải viết "small but durable" thay vì kết luận dứt khoát.',
    color_hex='F3E5F5')
P()

# ============================================================
# PHẦN 8 — KẾT LUẬN
# ============================================================
H2('8. Kết luận — Checklist cho đọc báo cáo tâm lý học')

P('Khi thầy/anh đọc báo cáo tâm lý học có số liệu thống kê, cần kiểm tra:', italic=True)
P()
P('[ ] 1. CI hay CrI? Xem Methods section có ghi "Bayesian" không.', size=11)
P('[ ] 2. MD hay SMD? Nếu nhiều NC dùng thang khác nhau → chắc chắn SMD (thường Hedges\' g).', size=11)
P('[ ] 3. Effect size có được báo cáo kèm CI/CrI không? (APA 7 yêu cầu)', size=11)
P('[ ] 4. Random-effects hay fixed-effects? (mặc định mong đợi random-effects).', size=11)
P('[ ] 5. Heterogeneity (I², τ²) được báo cáo không? Nếu I² > 75 % → pooled estimate có thể misleading.', size=11)
P('[ ] 6. Với Bayesian: prior được nêu rõ không? (weakly informative thì OK, strong informative cần diễn giải cẩn thận)', size=11)
P('[ ] 7. Diễn giải lâm sàng: áp Cohen\'s cutoffs (|SMD|<0,2 trivial...). CI/CrI nằm hoàn toàn trong SESOI (thường ±0,2) → effect quá nhỏ về mặt lâm sàng.', size=11)
P('[ ] 8. Tránh quên NNT (Number Needed to Treat) — rất hữu ích cho practical decision.', size=11)
P()

note_box(
    'TÓM TẮT CHO THẦY: (1) Trong tâm lý học lâm sàng hiện đại, CI và CrI ĐỀU được chấp nhận — '
    'không cần quá lo về "đúng về philosophy" vì kết luận lâm sàng gần như giống nhau. (2) Khi gặp '
    '"KTC 95 % CrI" trong báo cáo Bayesian MA → đọc giống như "CI thông thường" với ghi chú là '
    'kết quả từ Bayesian analysis. (3) Điểm CHÍNH YẾU cần chú ý là ĐỘ LỚN effect size (áp Cohen '
    'cutoffs) và KTC có chứa 0 hay không, chứ KHÔNG PHẢI chuyện CI vs CrI.',
    color_hex='E8F5E9')
P()

# ============================================================
# PHỤ LỤC
# ============================================================
P('─' * 70, color=RGBColor(0x99, 0x99, 0x99))
H3('Phụ lục A: Nguồn tham khảo')
P('• Kruschke JK. "Doing Bayesian Data Analysis" (2015, Academic Press) — chương 12: Credible Intervals',
  size=10)
P('• Cumming G. "Understanding The New Statistics: Effect Sizes, Confidence Intervals, and Meta-Analysis" (2012, Routledge)',
  size=10)
P('• American Psychological Association. "Publication Manual of the APA" (7th ed., 2020)',
  size=10)
P('• Lakens D. "Equivalence Tests: A Practical Primer for t Tests, Correlations, and Meta-Analyses" (2017, Soc Psychol Personal Sci)',
  size=10)
P('• Borenstein M et al. "Introduction to Meta-Analysis" (2nd ed., 2021, Wiley) — chương 4 về SMD/Hedges\'g',
  size=10)
P('• Cochrane Handbook for Systematic Reviews of Interventions v6.4 (2023) — chương 10 về effect measures',
  size=10)
P()

H3('Phụ lục B: Các bài Bayesian trong thư viện dự án Lo âu')
P('• QT029 — Li et al. 2025 (BMC Psychiatry) — Bayesian NMA 30 RCT điều trị lo âu trẻ em', size=10)
P('• QT039 — Xian, Zhang & Jiang 2024 (J Affect Disord) — Bayesian NMA 30 RCT cho SAD', size=10)
P('• QT049 — Zhang, Liang & Kang 2026 (J Clin Psychol) — Bayesian hierarchical MA 31 RCT CBT học đường', size=10)
P()

P('Tài liệu này do Nhóm dự án Lo âu soạn 15/04/2026. Phiên bản 2 — đã rà soát phê phán bởi người dịch '
  '+ bối cảnh hoá vào methodology tâm lý học. Mọi thắc mắc hoặc chỉnh sửa, vui lòng phản hồi qua dự án.',
  italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT):,} bytes')
