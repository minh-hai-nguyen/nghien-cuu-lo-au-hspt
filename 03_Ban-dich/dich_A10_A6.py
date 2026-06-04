# -*- coding: utf-8 -*-
"""Dịch A10 Hải Phòng (VN) + A6 Zheng MXH+Stress+Sleep (TQ)"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

# ===== A10 Hải Phòng — Bài VN → nhận xét + phản biện =====
doc1 = create_doc()
add_p(doc1, 'Link: https://doi.org/10.51403/0868-2836/2024/1571', size=10)
add_heading(doc1, 'Thực trạng sức khỏe tâm thần và một số yếu tố liên quan của học sinh tại 2 trường THPT huyện Vĩnh Bảo, Hải Phòng năm 2023', 1)
h = doc1.add_paragraph()
r = h.add_run('Mental Health Status and Related Factors among Students at 2 High Schools in Vinh Bao District, Hai Phong, 2023')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc1, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc1, [
    ('Tiêu đề', 'Thực trạng sức khỏe tâm thần và một số yếu tố liên quan của HS tại 2 trường THPT huyện Vĩnh Bảo, Hải Phòng năm 2023'),
    ('Tác giả', 'Phạm Thị Ngọc (1)*, Hoàng Thị Giang (1), Phạm Khánh Linh (2), Vũ Thị Châu (3)'),
    ('Cơ quan', '(1) ĐH Y Dược Hải Phòng; (2) BV Vinmec Phú Quốc; (3) TT Kiểm dịch Y tế QT Hải Phòng'),
    ('Tạp chí', 'Tạp chí Y học Dự phòng, Tập 34, Số 1 Phụ bản, 2024, tr. 115–123'),
    ('DOI', '10.51403/0868-2836/2024/1571'),
    ('Loại NC', 'Cắt ngang mô tả'),
    ('Mẫu', '420 HS lớp 10–12: THPT Cộng Hiền (210) + TT GDNN-GDTX Vĩnh Bảo (210); 10/2022–5/2023'),
    ('Công cụ', 'DASS-21 (trầm cảm + lo âu + căng thẳng)'),
])

add_heading(doc1, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'TC Y học Dự phòng — có DOI, peer-reviewed VN.',
    'So sánh 2 loại trường: THPT (Cộng Hiền) vs GDNN-GDTX (Vĩnh Bảo) — hiếm.',
    'n = 420 HS — đủ sức mạnh thống kê.',
    'DASS-21 — cùng công cụ với nhiều bài VN khác (VN15-19).',
    'Hải Phòng — thành phố lớn miền Bắc, khác Hà Nội (đô thị lớn) và Huế (miền Trung).',
    'Lo âu 53,81% (THPT) vs 43,33% (GDNN-GDTX) — THPT CAO HƠN.',
    'Yếu tố: áp lực học tập, bạn bè, bạo lực học đường, kỳ vọng gia đình, mâu thuẫn gia đình.',
]:
    add_p(doc1, f'• {b}')
add_p(doc1, 'Hạn chế:', bold=True)
for b in [
    'Chỉ 2 trường Vĩnh Bảo — không đại diện toàn Hải Phòng.',
    'TC Y học Dự phòng — không PubMed/Scopus.',
    'Bài chủ yếu trầm cảm — lo âu ít chi tiết hơn.',
]:
    add_p(doc1, f'• {b}')

add_heading(doc1, 'KẾT QUẢ CHÍNH', 2)
add_heading(doc1, 'Bảng 1. Tỷ lệ DASS-21 — So sánh 2 trường (n = 420)', 3)
add_table(doc1,
    ['Chỉ số', 'THPT Cộng Hiền', 'TT GDNN-GDTX', 'Ghi chú'],
    [['Trầm cảm', '42,38%', '32,38%', 'THPT cao hơn'],
     ['Lo âu', '53,81%', '43,33%', 'THPT cao hơn'],
     ['Căng thẳng', '49,05%', '28,57%', 'THPT CAO HƠN NHIỀU']],
    widths=[3.0, 3.0, 3.0, 3.5])

add_heading(doc1, 'Bảng 2. Yếu tố liên quan', 3)
add_table(doc1,
    ['Kết quả', 'Yếu tố nguy cơ', 'p'],
    [['Trầm cảm', 'Nữ giới; không hòa đồng bạn bè', '<0,05'],
     ['Lo âu', 'Áp lực học tập + thi cử CAO; không hòa đồng bạn bè; bạo lực học đường', '<0,05'],
     ['Căng thẳng', 'Trường THPT Cộng Hiền; áp lực kỳ vọng gia đình; mâu thuẫn gia đình', '<0,05']],
    widths=[2.5, 6.5, 1.5])
add_p(doc1, 'Lo âu liên quan: áp lực HT + bạo lực học đường + quan hệ bạn bè — phù hợp UNICEF VN22, Islam QT31 (bắt nạt AOR=1,68), Wen QT08 (áp lực OR=11,58).', size=9, italic=True)

add_p(doc1, 'Đối chiếu liên bài', bold=True)
add_p(doc1, 'Lo âu 53,81% (DASS-21 THPT) — phù hợp dải 40–86% DASS-21 VN (Hoa 40,6% GAD-7; Bảo Quyên 86,2%; Thảo Vi 65,8%; An Giang 61,2%). THPT cao hơn GDNN-GDTX — gợi ý áp lực thi cử THPT lớn hơn (phù hợp VN21: thi vào lớp 10 β=4,73). Bạo lực học đường = yếu tố lo âu — phù hợp UNICEF VN22 (bắt nạt nghiêm trọng), Islam QT31 (AOR=1,68). Mâu thuẫn gia đình — phù hợp Vĩnh Lộc VN20 (cha mẹ ly hôn), Vĩnh Long A9 (nghiện rượu, bạo lực).')

add_red_heading(doc1, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc1, 'Điểm mạnh:', bold=True)
for s in ['TC YHDP có DOI. n=420. So sánh 2 loại trường — hiếm.', 'DASS-21 — so sánh được với NC VN khác.', 'Hải Phòng miền Bắc — mở rộng vùng. Phù hợp UNICEF VN22 (5 tỉnh nhưng không có HP).', 'Yếu tố: bạo lực HĐ, áp lực, gia đình — phù hợp y văn.']:
    add_red(doc1, f'• {s}')
add_red(doc1, 'Hạn chế:', bold=True)
for s in ['Chỉ 2 trường Vĩnh Bảo — nông thôn HP, không đại diện TP.', 'TC VN không PubMed.', 'Bài tập trung trầm cảm — lo âu ít chi tiết (thiếu OR cụ thể).', 'DASS-21 sàng lọc — không chẩn đoán. So V-NAMHS 2,3%.']:
    add_red(doc1, f'• {s}')
add_red(doc1, 'Gap: NC lo âu riêng (GAD-7) tại HP. So sánh đô thị HP vs nông thôn Vĩnh Bảo. Thêm MXH, giấc ngủ.')

add_p(doc1, 'Đánh giá: ⭐⭐⭐ Trung bình. TC VN DOI, n=420, so sánh 2 trường, HP miền Bắc.', bold=True)

doc1.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), '45_HaiPhong_2024.docx'))

# ===== A6 Zheng MXH+Stress+Sleep → Anxiety (TQ) =====
doc2 = create_doc()
add_p(doc2, 'Link: https://doi.org/10.2147/PRBM.S522652', size=10)
add_heading(doc2, 'Tác động của nghiện mạng xã hội, căng thẳng học tập và chất lượng giấc ngủ lên triệu chứng lo âu: Nghiên cứu cắt ngang ở học sinh nghề Trung Quốc', 1)
h = doc2.add_paragraph()
r = h.add_run('The Effects of Social Media Addiction, Academic Stress, and Sleep Quality on Anxiety Symptoms: A Cross-Sectional Study of Chinese Vocational Students')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc2, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc2, [
    ('Tiêu đề gốc', 'The Effects of Social Media Addiction, Academic Stress, and Sleep Quality on Anxiety Symptoms: A Cross-Sectional Study of Chinese Vocational Students'),
    ('Tác giả', 'GuangFeng Zheng (1), HaoYan Peng (2)'),
    ('Cơ quan', '(1) Phòng Công tác SV, ĐH Bách khoa Cơ điện Quảng Đông; (2) ĐH Nghề Quảng Châu, TQ'),
    ('Tạp chí', 'Psychology Research and Behavior Management, 2025'),
    ('DOI', '10.2147/PRBM.S522652'),
    ('Loại NC', 'Cắt ngang — phân tích trung gian (mediation)'),
    ('Mẫu', '469 HS nghề 12–18 tuổi, Quảng Đông, TQ'),
    ('Công cụ', 'SMAS (6 mục, nghiện MXH) + Academic Stress Scale (16 mục) + PSQI (giấc ngủ) + GAD-7 (lo âu) + GSES (self-efficacy)'),
    ('Truy cập', 'Open Access — PMC'),
])
add_page_ref(doc2, '1–14', 'Psych Res Behav Mgmt', '2025')

add_heading(doc2, 'TÓM TẮT', 2)
p = doc2.add_paragraph()
r = p.add_run('Phương pháp: Phân tích tương quan + trung gian (mediation) trên 469 HS nghề 12–18 tuổi. Đo: SMAS (nghiện MXH), Academic Stress (áp lực HT), PSQI (giấc ngủ), GAD-7 (lo âu), GSES (self-efficacy — hiệu quả bản thân). Kiểm tra self-efficacy như biến TRUNG GIAN.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc2.add_paragraph()
r = p.add_run('Kết quả: Tương quan trực tiếp → lo âu: giấc ngủ kém r = 0,816 (MẠNH NHẤT); áp lực HT r = 0,505; MXH r = 0,415. Hồi quy: giấc ngủ β = 0,615 (mạnh nhất); áp lực β = 0,223; MXH β = 0,153. Self-efficacy là TRUNG GIAN: MXH gián tiếp 63,13% tổng tác động; áp lực 55,84%; giấc ngủ 24,63%.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0, 0)

add_heading(doc2, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Psych Res Behav Mgmt — PMC indexed. Open Access.',
    'MÔ HÌNH TAM GIÁC: MXH + Áp lực HT + Giấc ngủ → Lo âu (GAD-7). Ít NC nào đo CẢ 3 cùng lúc.',
    'Giấc ngủ β = 0,615 — MẠNH NHẤT. Phù hợp Zhu 2025 (QT05: <5h AOR = 13,71).',
    'Self-efficacy là TRUNG GIAN — phát hiện CƠ CHẾ. MXH → giảm self-efficacy → tăng lo âu.',
    'MXH gián tiếp 63,13% — MXH ảnh hưởng lo âu CHỦ YẾU QUA self-efficacy (không trực tiếp).',
    'GAD-7 — cùng công cụ với Hoa 2024 (VN01) và A8 VN COVID.',
    'HS nghề 12–18 — nhóm ÍT NC (đa số NC tập trung THPT chung).',
]:
    add_p(doc2, f'• {b}')
add_p(doc2, 'Hạn chế:', bold=True)
for b in [
    'Chỉ Quảng Đông TQ — HS nghề (khác THPT chung VN).',
    'n = 469 — trung bình.',
    'Cắt ngang — mediation GIẢ ĐỊNH nhân quả.',
    'SMAS chỉ 6 mục — đo sơ lược.',
]:
    add_p(doc2, f'• {b}')

add_heading(doc2, 'KẾT QUẢ CHI TIẾT', 2)
add_heading(doc2, 'Bảng 1. Tương quan và hồi quy — 3 yếu tố → Lo âu', 3)
add_table(doc2,
    ['Yếu tố', 'Tương quan r', 'Hồi quy β', 'Trung gian gián tiếp', 'Ý nghĩa'],
    [['Giấc ngủ kém (PSQI)', 'r = 0,816***', 'β = 0,615***', '24,63% qua self-efficacy', 'MẠNH NHẤT — trực tiếp chủ yếu'],
     ['Áp lực HT', 'r = 0,505***', 'β = 0,223***', '55,84% qua self-efficacy', 'Trung gian MẠNH'],
     ['Nghiện MXH (SMAS)', 'r = 0,415***', 'β = 0,153***', '63,13% qua self-efficacy', 'Chủ yếu GIÁ TIẾP — qua self-efficacy']],
    widths=[3.0, 2.0, 2.0, 3.5, 3.5])
add_p(doc2, 'MXH ảnh hưởng lo âu CHỦ YẾU QUA self-efficacy (63% gián tiếp). Giấc ngủ ảnh hưởng TRỰC TIẾP (75% trực tiếp). Áp lực HT cân bằng (56% gián tiếp). Self-efficacy = "cổng" trung gian.', size=9, italic=True)

add_heading(doc2, 'Bảng 2. So sánh với NC khác', 3)
add_table(doc2,
    ['NC', 'Yếu tố', 'Kết quả', 'Phù hợp?'],
    [['Zheng 2025 (bài này)', 'Giấc ngủ', 'β = 0,615 (mạnh nhất)', 'PHÙ HỢP'],
     ['Zhu 2025 (QT05)', 'Giấc ngủ <5h', 'AOR = 13,71', 'PHÙ HỢP — giấc ngủ mạnh nhất'],
     ['Norway 2025 (QT21)', 'MXH', 'Giải thích MỘT PHẦN xu hướng', 'PHÙ HỢP — MXH không chính'],
     ['Nature 2025 (QT27)', 'MXH', 'g = 0,46 (thời gian)', 'PHÙ HỢP — VTN SKTT dùng nhiều hơn'],
     ['Li 2025 (QT22)', 'Screen time', 'Dọc YẾU (p=0,443 lo âu)', 'PHÙ HỢP — hai chiều'],
     ['Wen 2020 (QT08)', 'Áp lực HT', 'OR = 11,58', 'PHÙ HỢP — áp lực mạnh'],
     ['VN19 Thảo Vi', 'Lạc quan', 'β gián tiếp = −0,24', 'TƯƠNG TỰ — trung gian tâm lý']],
    widths=[3.0, 2.0, 3.0, 3.5])

add_heading(doc2, 'THẢO LUẬN VÀ KẾT LUẬN', 2)
add_p(doc2, 'Mô hình tam giác MXH + Áp lực + Giấc ngủ → Lo âu, với self-efficacy là trung gian. Phát hiện quan trọng: MXH ảnh hưởng lo âu chủ yếu GIÁN TIẾP qua giảm self-efficacy (63%). Nghĩa là: MXH không trực tiếp "gây" lo âu — mà MXH → giảm tự tin → tăng lo âu. Gợi ý can thiệp: TĂNG SELF-EFFICACY (tự tin, kỹ năng) có thể CHẶN tác động MXH → lo âu.')
add_p(doc2, 'Giấc ngủ kém → lo âu MẠNH NHẤT (β=0,615) — phù hợp Zhu 2025 (AOR=13,71). Can thiệp giấc ngủ có thể HIỆU QUẢ nhất.')
add_p(doc2, 'Áp lực HT → lo âu (β=0,223), 56% qua self-efficacy. Phù hợp Wen 2020 (OR=11,58), VN21 (β=4,73), Norway 2025 (bất mãn trường giải thích chính).')

add_heading(doc2, 'TÀI LIỆU THAM KHẢO', 2)
for ref in [
    'Zheng, G.F. & Peng, H.Y. (2025). Effects of Social Media Addiction, Academic Stress, and Sleep Quality on Anxiety. Psych Res Behav Mgmt.',
    '(Xem đầy đủ 50+ TLTK trong bài gốc)',
]:
    add_p(doc2, ref, size=10)

add_abbreviation_table(doc2, [
    ('SMAS', 'Social Media Addiction Scale — Thang Nghiện MXH'),
    ('PSQI', 'Pittsburgh Sleep Quality Index — Chỉ số Chất lượng Giấc ngủ Pittsburgh'),
    ('GAD-7', 'Generalized Anxiety Disorder 7-item Scale'),
    ('GSES', 'General Self-Efficacy Scale — Thang Hiệu quả Bản thân Tổng quát'),
    ('Self-efficacy', 'Hiệu quả bản thân — niềm tin vào khả năng hoàn thành nhiệm vụ'),
])

add_red_heading(doc2, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc2, 'Điểm mạnh:', bold=True)
for s in [
    'MÔ HÌNH TAM GIÁC: MXH + Áp lực + Giấc ngủ → Lo âu — ít NC nào đo CẢ 3.',
    'Self-efficacy là TRUNG GIAN — phát hiện CƠ CHẾ. Can thiệp tăng self-efficacy có cơ sở.',
    'GAD-7 — so sánh được với Hoa 2024 (VN01) và A8 VN COVID.',
    'MXH 63% gián tiếp — giải thích tại sao Li 2025 (QT22) tìm thấy dọc YẾU: MXH tác động qua trung gian.',
    'Open Access PMC.',
]:
    add_red(doc2, f'• {s}')
add_red(doc2, 'Hạn chế:', bold=True)
for s in [
    'Chỉ TQ Quảng Đông — HS nghề (khác THPT VN).',
    'n = 469 — trung bình. Cắt ngang — mediation giả định nhân quả.',
    'SMAS 6 mục — đo sơ lược nghiện MXH.',
    'Không đo cyberbullying, so sánh xã hội — yếu tố MXH quan trọng (Nature QT27).',
]:
    add_red(doc2, f'• {s}')
add_red(doc2, 'Gap: Mô hình tam giác tại VN — MXH + áp lực + giấc ngủ + self-efficacy → lo âu HS THPT VN. So sánh THPT vs THCS. Can thiệp: tăng self-efficacy tại trường + quản lý giấc ngủ.')

add_p(doc2, 'Đánh giá: ⭐⭐⭐⭐ Cao. Mô hình tam giác + trung gian, GAD-7, phát hiện cơ chế self-efficacy.', bold=True)

doc2.save(os.path.join(os.path.dirname(os.path.abspath(__file__)), '46_Zheng_MXH_Stress_Sleep_2025.docx'))

# Verify
import docx as dx
for name, fp in [('A10', '45_HaiPhong_2024.docx'), ('A6', '46_Zheng_MXH_Stress_Sleep_2025.docx')]:
    d = dx.Document(os.path.join(os.path.dirname(os.path.abspath(__file__)), fp))
    t = '\n'.join([p.text for p in d.paragraphs])
    print(f'{name}: {len(t)} chars, {len(d.tables)} tables')
