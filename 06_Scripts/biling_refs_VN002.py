# -*- coding: utf-8 -*-
"""
Convert VN002 FULL References list to bilingual:
  Line 1: English original (kept as-is)
  Line 2: Vietnamese translation (italic, smaller)
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')

# Bilingual pairs: EN (as currently in docx) → VN translation
BILING = [
    ('American Psychiatric Association. 2013. "Diagnostic and Statistical Manual of Mental Disorders: Fifth Edition, DSM-5." Washington DC: American Psychiatric Publishing.',
     'Hiệp hội Tâm thần Hoa Kỳ. 2013. "Sổ tay Chẩn đoán và Thống kê Rối loạn Tâm thần: Tái bản thứ Năm, DSM-5." Washington DC: Nhà xuất bản Hiệp hội Tâm thần Hoa Kỳ.'),

    ('Bitsko, R., H.R. Adams, J. Holbrook, A. Vierhile, P. Morrison and P.W. Fisher. 2019. "2.50 Diagnostic Interview Schedule for Children, Version 5 (DISC-5): Development and Validation of ADHD and Tic Disorder Modules." Journal of the American Academy of Child & Adolescent Psychiatry 58(10): Supplement.',
     'Bitsko R. và cộng sự. 2019. "2.50 Bảng Phỏng vấn Chẩn đoán cho Trẻ em, Phiên bản 5 (DISC-5): Phát triển và Thẩm định các Module ADHD và Rối loạn Tic." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 58(10): Phụ lục.'),

    ('Blum, R., M. Sudhinaraset and M.R. Emerson. 2012. "Youth at Risk: Suicidal Thoughts and Attempts in Vietnam, China, and Taiwan." Journal of Adolescent Health 50(3):S37–S44. doi: 10.1016/j.jadohealth.2011.12.006',
     'Blum R., Sudhinaraset M. và Emerson M.R. 2012. "Vị thành niên có nguy cơ: Ý nghĩ và Hành vi Tự sát tại Việt Nam, Trung Quốc và Đài Loan." Tạp chí Sức khoẻ Vị thành niên 50(3):S37–S44.'),

    ('Bộ Giáo dục và Đào tạo. 2022a. "Quyết định số 2138/QĐ-BGDĐT ngày 03 tháng 08 năm 2022 ban hành Kế hoạch giáo dục sức khỏe tâm thần cho trẻ em, học sinh giai đoạn 2022-2025." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch. Đã bản địa hoá trước khi báo cáo xuất bản.)'),

    ('Bộ Giáo dục và Đào tạo. 2022b. "Quyết định số 1442/QĐ-BGDĐT ngày 01 tháng 6 năm 2022 ban hành Chương trình giáo dục sức khỏe tâm thần cho trẻ em, học sinh giai đoạn 2022-2025 Của Ngành Giáo Dục." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Bộ Y tế. 2022. "Công văn số 2213/BYT-DP ngày 29 tháng 4 năm 2022 về việc Tạm dừng áp dụng khai báo y tế nội địa." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Canino, G. and M. Alegria. 2008. "Psychiatric Diagnosis — Is It Universal or Relative to Culture?" Journal of Child Psychology and Psychiatry 49(3):237–50.',
     'Canino G. và Alegria M. 2008. "Chẩn đoán tâm thần — Có tính phổ quát hay tuỳ thuộc văn hoá?" Tạp chí Tâm lý học và Tâm thần học Trẻ em 49(3):237–50.'),

    ('Cao Tiến Đức. 2020. "Sức khoẻ tâm thần: thực trạng, thách thức và những tiến bộ mới trong chẩn đoán và điều trị." Tạp chí Nội khoa Việt Nam 19:15–20.',
     '(Tác phẩm tiếng Việt gốc — không cần dịch.)'),

    ('Chính phủ. 2022. "Nghị quyết 38/NQ-CP ngày 17 tháng 3 năm 2022 ban hành Chương trình phòng, chống dịch Covid-19." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Colizzi, M., A. Lasalvia and M. Ruggeri. 2020. "Prevention and Early Intervention in Youth Mental Health: Is It Time for a Multidisciplinary and Trans-Diagnostic Model for Care?" International Journal of Mental Health Systems 14.',
     'Colizzi M., Lasalvia A. và Ruggeri M. 2020. "Phòng ngừa và Can thiệp Sớm trong Sức khoẻ Tâm thần Thanh thiếu niên: Đã đến lúc cho Mô hình Chăm sóc Đa ngành và Xuyên Chẩn đoán?" Tạp chí Quốc tế về Hệ thống Sức khoẻ Tâm thần 14.'),

    ('COVID-19 Mental Disorders Collaborators. 2021. "Global Prevalence and Burden of Depressive and Anxiety Disorders in 204 Countries and Territories in 2020 Due to the Covid-19 Pandemic." The Lancet 398(10312):1700–12.',
     'Nhóm Cộng tác Rối loạn Tâm thần COVID-19. 2021. "Tỷ lệ và Gánh nặng Toàn cầu của Rối loạn Trầm cảm và Lo âu tại 204 Quốc gia và Vùng lãnh thổ năm 2020 do Đại dịch COVID-19." The Lancet 398(10312):1700–12.'),

    ('Demyttenaere, K. et al. 2004. "Prevalence, Severity, and Unmet Need for Treatment of Mental Disorders in the WHO World Mental Health Surveys." JAMA 291(21):2581–90.',
     'Demyttenaere K. và cộng sự. 2004. "Tỷ lệ, Mức độ nghiêm trọng và Nhu cầu Điều trị chưa được Đáp ứng đối với Rối loạn Tâm thần trong các Khảo sát Sức khoẻ Tâm thần Thế giới của WHO." JAMA 291(21):2581–90.'),

    ('Erskine, H.E. et al. 2015. "A Heavy Burden on Young Minds: The Global Burden of Mental and Substance Use Disorders in Children and Youth." Psychological Medicine 45(7):1551–63.',
     'Erskine H.E. và cộng sự. 2015. "Một Gánh nặng cho Tâm trí Trẻ: Gánh nặng Toàn cầu của Rối loạn Tâm thần và Sử dụng Chất ở Trẻ em và Thanh niên." Y học Tâm lý 45(7):1551–63.'),

    ('Erskine, H.E. et al. 2016. "Long-Term Outcomes of Attention-Deficit/Hyperactivity Disorder and Conduct Disorder: A Systematic Review and Meta-Analysis." Journal of the American Academy of Child & Adolescent Psychiatry 55(10):841–50.',
     'Erskine H.E. và cộng sự. 2016. "Kết quả dài hạn của Rối loạn Tăng động Giảm Chú ý và Rối loạn Hành vi: Tổng quan Hệ thống và Phân tích Meta." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 55(10):841–50.'),

    ('Erskine, H.E., S. Blondell, M. Enright, J. Shadid, Y. Wado, F.M. Wekesah, A.E. Wahdi, S.A. Wilopo, L.M. Vu, H.T.K. Dao, V.D. Nguyen, M.R. Emerson, S.L. Fine, Mengmeng Li, R.W. Blum, H.A. Whiteford and J.G. Scott. 2021. "Measuring the Prevalence of Mental Disorders in Adolescents in Kenya, Indonesia, and Vietnam: Study Protocol for the National Adolescent Mental Health Surveys." Journal of Adolescent Health.',
     'Erskine H.E. và cộng sự (bao gồm Vũ Mạnh Lợi, Đào Thị Khánh Hoa, Nguyễn Đức Vinh từ Việt Nam). 2021. "Đo lường Tỷ lệ Rối loạn Tâm thần ở Vị thành niên tại Kenya, Indonesia và Việt Nam: Đề cương Nghiên cứu cho các Khảo sát Sức khoẻ Tâm thần Vị thành niên Quốc gia." Tạp chí Sức khoẻ Vị thành niên.'),

    ('Ferrari, A.J. et al. 2013. "The Epidemiological Modelling of Major Depressive Disorder: Application for the Global Burden of Disease Study 2010." PLOS One 8(7): e69637.',
     'Ferrari A.J. và cộng sự. 2013. "Mô hình hoá Dịch tễ học Rối loạn Trầm cảm Nặng: Ứng dụng cho Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2010." PLOS One 8(7): e69637.'),

    ('GBD 2019 Mental Disorders Collaborators. 2022. "Global, Regional, and National Burden of 12 Mental Disorders in 204 Countries and Territories, 1990–2019: A Systematic Analysis for the Global Burden of Disease Study 2019." The Lancet Psychiatry 9(2):137–50.',
     'Nhóm Cộng tác GBD 2019 Rối loạn Tâm thần. 2022. "Gánh nặng Toàn cầu, Khu vực và Quốc gia của 12 Rối loạn Tâm thần tại 204 Quốc gia và Vùng lãnh thổ, 1990–2019: Phân tích Hệ thống cho Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019." The Lancet Psychiatry 9(2):137–50.'),

    ('General Statistical Office. 2020. Kết quả toàn bộ Tổng điều tra dân số và nhà ở năm 2019 / Completed Results of the 2019 Viet Nam Population and Housing Census. Hà Nội: NXB Thống kê.',
     '(Song ngữ gốc — không cần dịch bổ sung. Tác giả: Tổng cục Thống kê Việt Nam.)'),

    ('Hafekost, J. et al. 2016. "Methodology of Young Minds Matter: The Second Australian Child and Adolescent Survey of Mental Health and Wellbeing." Australian and New Zealand Journal of Psychiatry 50(9):866–75.',
     'Hafekost J. và cộng sự. 2016. "Phương pháp của Young Minds Matter: Khảo sát Sức khoẻ Tâm thần và Hạnh phúc Trẻ em và Vị thành niên Úc lần thứ Hai." Tạp chí Tâm thần học Úc và New Zealand 50(9):866–75.'),

    ('Hafstad, G.S. and E. Augusti. 2021. "A Lost Generation? Covid-19 and Adolescent Mental Health." The Lancet Psychiatry 8(8):640–41.',
     'Hafstad G.S. và Augusti E. 2021. "Một Thế hệ Mất mát? COVID-19 và Sức khoẻ Tâm thần Vị thành niên." The Lancet Psychiatry 8(8):640–41.'),

    ('Hoang M.D., T.T. Lam, A. Dao and B. Weiss. 2020. "Mental Health Literacy at the Public Health Level in Low and Middle Income Countries: An Exploratory Mixed Methods Study in Vietnam." PLOS One.',
     'Hoàng Minh Đặng, Lâm T.T., Đào A. và Weiss B. 2020. "Kiến thức về Sức khoẻ Tâm thần ở Cấp độ Y tế Công cộng tại các Quốc gia Thu nhập Thấp và Trung bình: Nghiên cứu Hỗn hợp Thăm dò tại Việt Nam." PLOS One.'),

    ('Hoang M.D., B. Weiss, T. Lam and H. Ho. 2018. "Mental Health Literacy and Intervention Program Adaptation in the Internationalization of School Psychology for Vietnam." Psychology in the School 55(8):941–54.',
     'Hoàng Minh Đặng, Weiss B., Lâm T. và Hồ H. 2018. "Kiến thức Sức khoẻ Tâm thần và Thích ứng Chương trình Can thiệp trong Quốc tế hoá Tâm lý học Trường học cho Việt Nam." Tâm lý học trong Trường học 55(8):941–54.'),

    ('Hoàng Thế Hải and Bùi Thị Thanh Diệu. 2021. "Thái độ kỳ thị đối với người mắc bệnh tâm thần của học sinh trung học phổ thông." Tạp chí Khoa học và Công nghệ 19(2):33–37.',
     '(Tác phẩm tiếng Việt gốc — không cần dịch.)'),

    ('Jones E.A., A.K. Mitra and A.R. Bhuiyan. 2021. "Impact of Covid-19 on Mental Health in Adolescents: A Systematic Review." International Journal of Environmental Research and Public Health 18(5):2470.',
     'Jones E.A., Mitra A.K. và Bhuiyan A.R. 2021. "Tác động của COVID-19 lên Sức khoẻ Tâm thần ở Vị thành niên: Tổng quan Hệ thống." Tạp chí Quốc tế về Nghiên cứu Môi trường và Y tế Công cộng 18(5):2470.'),

    ('Kamimura, A. et al. 2018. "Perceptions of Mental Health and Mental Health Services among College Students in Vietnam and the United States." Asian Journal of Psychiatry 37:15–19.',
     'Kamimura A. và cộng sự. 2018. "Nhận thức về Sức khoẻ Tâm thần và Dịch vụ Sức khoẻ Tâm thần ở Sinh viên Đại học tại Việt Nam và Hoa Kỳ." Tạp chí Tâm thần học Châu Á 37:15–19.'),

    ('Keynejad, R.C., J. Spagnolo and G. Thornicroft. 2022. "Mental Healthcare in Primary and Community-Based Settings: Evidence Beyond the WHO Mental Health Gap Action Programme (mhGAP) Intervention Guide." Evidence-Based Mental Health.',
     'Keynejad R.C., Spagnolo J. và Thornicroft G. 2022. "Chăm sóc Sức khoẻ Tâm thần tại Bối cảnh Y tế Ban đầu và Cộng đồng: Bằng chứng Vượt ra ngoài Hướng dẫn Can thiệp Chương trình Hành động Lấp khoảng trống Sức khoẻ Tâm thần của WHO (mhGAP)." Sức khoẻ Tâm thần Dựa trên Bằng chứng.'),

    ('Kim, J.H.J., W. Tsai, T. Kodish, L.T. Trung, A.S. Lau and B. Weiss. 2019. "Cultural Variation in Temporal Associations among Somatic Complaints, Anxiety, and Depressive Symptoms in Adolescence." Journal of Psychosomatic Research 124:109763.',
     'Kim J.H.J., Tsai W., Kodish T., Lê T. Trung, Lau A.S. và Weiss B. 2019. "Khác biệt Văn hoá trong Mối liên hệ Thời gian giữa Khiếu nại Cơ thể, Triệu chứng Lo âu và Trầm cảm ở Vị thành niên." Tạp chí Nghiên cứu Tâm lý học Cơ thể 124:109763.'),

    ('Mai D., N.N.K. Pham, S. Wallick and B.K. Nastasi. 2014. "Perceptions of Mental Illness and Related Stigma among Vietnamese Populations: Findings from a Mixed Method Study." Journal of Immigrant and Minority Health 16:1294–98.',
     'Mai Đỗ, Phạm N.N.K., Wallick S. và Nastasi B.K. 2014. "Nhận thức về Bệnh Tâm thần và Kỳ thị Liên quan trong Quần thể người Việt: Phát hiện từ Nghiên cứu Phương pháp Hỗn hợp." Tạp chí Sức khoẻ Người nhập cư và Dân tộc Thiểu số 16:1294–98.'),

    ('Mckelvey, R.S., L.V. Baldassar, D.L. Sang and L. Roberts. 1999. "Vietnamese Parental Perceptions of Child and Adolescent Mental Illness." Journal of the American Academy of Child & Adolescent Psychiatry 38(10):1302–09.',
     'McKelvey R.S., Baldassar L.V., Đỗ L. Sang và Roberts L. 1999. "Nhận thức của Phụ huynh Việt Nam về Bệnh Tâm thần ở Trẻ em và Vị thành niên." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 38(10):1302–09.'),

    ('MOLISA (Bộ Lao động-Thương binh và Xã hội). 2012. Quyết định số 1364/QĐ-LĐTBXH ngày 2 tháng 10 năm 2012.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Nga L.L., I. Shochet, T. Tran, J. Fisher, A. Wurfl, N. Nguyen, J. Orr, R. Stocker and H. Nguyen. 2022. "Adaptation of a School-Based Mental Health Program for Adolescents in Vietnam." PLOS One 17(8): e0271959.',
     'La Linh Nga, Shochet I., Trần T., Fisher J., Wurfl A., Nguyễn N., Orr J., Stocker R. và Nguyễn H. 2022. "Thích ứng Chương trình Sức khoẻ Tâm thần Trường học cho Vị thành niên tại Việt Nam." PLOS One 17(8): e0271959.'),

    ('Nguyen, T., T. Tran, H. Tran, T. Tran and J. Fisher. 2019. "Challenges in Integrating Mental Health into Primary Care in Vietnam." in Innovation in Global Mental Health, edited by S. Okpaku: Springer, Cham.',
     'Nguyễn T., Trần T., Trần H., Trần T. và Fisher J. 2019. "Thách thức trong Lồng ghép Sức khoẻ Tâm thần vào Chăm sóc Y tế Ban đầu tại Việt Nam." trong Đổi mới Sức khoẻ Tâm thần Toàn cầu, do S. Okpaku biên tập: Springer, Cham.'),

    ('Nguyen, T.Q.C and T.H. Nguyen. 2018. "Mental Health Literacy: Knowledge of Depression among Undergraduate Students in Hanoi, Vietnam." International Journal of Mental Health Systems 24(12):19.',
     'Nguyễn T.Q.C và Nguyễn T.H. 2018. "Kiến thức Sức khoẻ Tâm thần: Hiểu biết về Trầm cảm ở Sinh viên Đại học tại Hà Nội, Việt Nam." Tạp chí Quốc tế về Hệ thống Sức khoẻ Tâm thần 24(12):19.'),

    ('Ormel, J. et al. 2017. "Functional Outcomes of Child and Adolescent Mental Disorders. Current Disorder Most Important but Psychiatric History Matters as Well." Psychological Medicine 47(7):1271–82.',
     'Ormel J. và cộng sự. 2017. "Kết quả Chức năng của Rối loạn Tâm thần ở Trẻ em và Vị thành niên. Rối loạn Hiện tại Quan trọng nhất nhưng Lịch sử Tâm thần cũng Quan trọng." Y học Tâm lý 47(7):1271–82.'),

    ('Pagliaro, C., M. Pearl, D. Lawrence, J.G. Scott and S. Diminic. 2021. "Estimating Demand for Mental Health Care among Australian Children and Adolescents: Findings from the Young Minds Matter Survey." Australian and New Zealand Journal of Psychiatry 56(11):1443–54.',
     'Pagliaro C., Pearl M., Lawrence D., Scott J.G. và Diminic S. 2021. "Ước tính Nhu cầu Chăm sóc Sức khoẻ Tâm thần ở Trẻ em và Vị thành niên Úc: Phát hiện từ Khảo sát Young Minds Matter." Tạp chí Tâm thần học Úc và New Zealand 56(11):1443–54.'),

    ('Patton G.C. et al. 2016. "Our Future: A Lancet Commission on Adolescent Health and Wellbeing." The Lancet 387(10036):2423–78.',
     'Patton G.C. và cộng sự. 2016. "Tương lai của Chúng ta: Uỷ ban The Lancet về Sức khoẻ và Hạnh phúc Vị thành niên." The Lancet 387(10036):2423–78.'),

    ('Person, S., C. Hagquist and D. Michelson. 2017. "Young Voices in Mental Health Care: Exploring Children\'s and Adolescents\' Service Experiences and Preferences." Clinical Child Psychology and Psychiatry 22(1):140–51.',
     'Person S., Hagquist C. và Michelson D. 2017. "Tiếng nói Trẻ trong Chăm sóc Sức khoẻ Tâm thần: Khám phá Trải nghiệm và Ưu tiên Dịch vụ của Trẻ em và Vị thành niên." Tâm lý học và Tâm thần học Trẻ em Lâm sàng 22(1):140–51.'),

    ('Samuels, F., N. Jones, T. Gupta, B.T. Dang and D.H. Le. 2018. Mental Health and Psychosocial Wellbeing among Children and Young People in Selected Provinces and Cities in Viet Nam. Hanoi: UNICEF.',
     'Samuels F., Jones N., Gupta T., Đặng B.T. và Lê D.H. 2018. Sức khoẻ Tâm thần và Hạnh phúc Tâm lý Xã hội ở Trẻ em và Người trẻ tại các Tỉnh và Thành phố được Chọn tại Việt Nam. Hà Nội: UNICEF.'),

    ('Schnyder, N., D. Lawrence, R. Panczak, M.G. Sawyer, H.A. Whiteford, P.M. Burgess and M.G. Harris. 2019. "Perceived need and barriers to adolescent mental health care: agreement between adolescents and their parents." Epidemiology and Psychiatric Sciences 29(e60).',
     'Schnyder N., Lawrence D., Panczak R., Sawyer M.G., Whiteford H.A., Burgess P.M. và Harris M.G. 2019. "Nhu cầu được cảm nhận và rào cản đối với chăm sóc sức khoẻ tâm thần vị thành niên: sự đồng ý giữa vị thành niên và phụ huynh." Dịch tễ học và Khoa học Tâm thần 29(e60).'),

    ('Shaffer, D., P. Fisher, C.P. Lucas, M.K. Dulcan and M.E.S. Stone. 2000. "NIMH Diagnostic Interview Schedule for Children Version IV (NIMH DISC-IV): Description, Differences from Previous Versions, and Reliability of Some Common Diagnoses." Journal of the American Academy of Child & Adolescent Psychiatry 39(1):28–38.',
     'Shaffer D., Fisher P., Lucas C.P., Dulcan M.K. và Stone M.E.S. 2000. "Bảng Phỏng vấn Chẩn đoán NIMH cho Trẻ em Phiên bản IV (NIMH DISC-IV): Mô tả, Khác biệt so với Phiên bản Trước, và Độ tin cậy của một số Chẩn đoán Phổ biến." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 39(1):28–38.'),

    ('Thủ tướng Chính phủ. 2011. "Quyết định số 1215/QĐ-TTg ngày 22/7/2011 phê duyệt Đề án trợ giúp xã hội và phục hồi chức năng cho người tâm thần giai đoạn 2011–2020." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Thủ tướng Chính phủ. 2020. "Quyết định 1929/QĐ-TTg ngày 25/11/2020 phê duyệt Chương trình trợ giúp xã hội và phục hồi chức năng cho người tâm thần, trẻ em tự kỷ và người rối nhiễu tâm trí dựa vào cộng đồng giai đoạn 2021–2030." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Thủ tướng Chính phủ. 2022. "Quyết định 155/QĐ-TTg ngày 29/01/2022 phê duyệt Kế hoạch quốc gia phòng chống bệnh không lây nhiễm và rối loạn sức khỏe tâm thần giai đoạn 2022–2025." Hà Nội.',
     '(Văn bản tiếng Việt gốc — không cần dịch.)'),

    ('Tran, T., H.T. Nguyen, I. Shochet, A. Wurfl, J. Orr, N. Nguyen, N. La, H. Nguyen, R. Stocker, T. Nguyen, M. Le and J. Fisher. 2020. "School-Based, Two-Arm, Parallel, Controlled Trial of a Culturally Adapted Resilience Intervention to Improve Adolescent Mental Health in Vietnam: Study Protocol." BMJ Open 10(10).',
     'Trần T., Nguyễn H.T., Shochet I., Wurfl A., Orr J., Nguyễn N., La N., Nguyễn H., Stocker R., Nguyễn T., Lê M. và Fisher J. 2020. "Thử nghiệm Đối chứng Song song Hai nhánh tại Trường học về Can thiệp Bền bỉ Thích ứng Văn hoá để Cải thiện Sức khoẻ Tâm thần Vị thành niên tại Việt Nam: Đề cương Nghiên cứu." BMJ Open 10(10).'),

    ('Truc, T.T., N.L.L.T. Vu and H.H.T. Bui. 2020. "Mental Health Literacy and Help-Seeking Preferences in High School Students in Ho Chi Minh City, Vietnam." School Mental Health 12:378–87.',
     'Thái T. Trúc, Vũ N.L.L.T. và Bùi H.H.T. 2020. "Kiến thức Sức khoẻ Tâm thần và Ưu tiên Tìm kiếm Trợ giúp ở Học sinh Trung học Phổ thông tại TP Hồ Chí Minh, Việt Nam." Sức khoẻ Tâm thần Trường học 12:378–87.'),

    ('UNICEF. 2020a. "Rapid Assessment of Social and Economic Impacts of Covid-19 on Children and Families in Viet Nam." Ha Noi: UNICEF.',
     'UNICEF. 2020a. "Đánh giá Nhanh Tác động Xã hội và Kinh tế của COVID-19 lên Trẻ em và Gia đình tại Việt Nam." Hà Nội: UNICEF.'),

    ('UNICEF. 2020b. "The Impact of Covid-19 on the Mental Health of Adolescents and Youth."',
     'UNICEF. 2020b. "Tác động của COVID-19 lên Sức khoẻ Tâm thần của Vị thành niên và Thanh niên."'),

    ('United Nations Children\'s Funds (UNICEF) and Overseas Development Institute (ODI). 2018. "The Nature of Suicide Amongst Children and Young People in Selected Provinces and Cities in Viet Nam." Hanoi: UNICEF.',
     'Quỹ Nhi đồng Liên Hợp Quốc (UNICEF) và Viện Phát triển Hải ngoại (ODI). 2018. "Bản chất của Tự sát ở Trẻ em và Người trẻ tại các Tỉnh và Thành phố được Chọn tại Việt Nam." Hà Nội: UNICEF.'),

    ('Wasserman, D., H.T.T. Tran, D.T.M. Pham, M. Goldstein, A. Nordenskiöld and C. Wasserman. 2008. "Suicidal Process, Suicidal Communication and Psychosocial Situation of Young Suicide Attempters in a Rural Vietnamese Community." World Psychiatry 7(1):47–53.',
     'Wasserman D., Trần H.T.T., Phạm D.T.M., Goldstein M., Nordenskiöld A. và Wasserman C. 2008. "Quá trình Tự sát, Giao tiếp về Tự sát và Tình huống Tâm lý Xã hội của Người trẻ Toan Tự sát trong một Cộng đồng Nông thôn Việt Nam." Tâm thần học Thế giới 7(1):47–53.'),

    ('Weiss, B., M. Dang, L. Trung, M.C. Nguyen, H.T.T. Nguyen and A. Pollack. 2014. "A Nationally-Representative Epidemiological and Risk Factor Assessment of Child Mental Health in Vietnam." International Perspectives in Psychology 3(3):139–53.',
     'Weiss B., Đặng M., Lê Trung, Nguyễn M.C., Nguyễn H.T.T. và Pollack A. 2014. "Đánh giá Dịch tễ học và Yếu tố Nguy cơ Đại diện Quốc gia về Sức khoẻ Tâm thần Trẻ em tại Việt Nam." International Perspectives in Psychology 3(3):139–53.'),

    ('World Health Organization. 2014. "Health for the World\'s Adolescents: A Second Chance in the Second Decade." Geneva.',
     'Tổ chức Y tế Thế giới. 2014. "Sức khoẻ cho Vị thành niên Thế giới: Cơ hội Thứ Hai trong Thập kỷ Thứ Hai." Geneva.'),

    ('World Health Organization. 2016. mhGAP Intervention Guide for Mental, Neurological and Substance Use Disorders in Non-Specialized Health Settings: Mental Health Gap Action Programme (mhGAP) – Version 2.0. Geneva.',
     'Tổ chức Y tế Thế giới. 2016. Hướng dẫn Can thiệp mhGAP cho Rối loạn Tâm thần, Thần kinh và Sử dụng Chất trong các Cơ sở Y tế Không Chuyên khoa: Chương trình Hành động Lấp khoảng trống Sức khoẻ Tâm thần (mhGAP) – Phiên bản 2.0. Geneva.'),

    ('World Health Organization, Ministry of Health and Ministry of Education and Training. 2022. The 2019 Global School-Based Student Health Survey in Viet Nam: Report. Manila: WHO Regional Office for the Western Pacific.',
     'Tổ chức Y tế Thế giới, Bộ Y tế và Bộ Giáo dục & Đào tạo. 2022. Khảo sát Sức khoẻ Học sinh Toàn cầu tại Trường học tại Việt Nam 2019: Báo cáo. Manila: Văn phòng Khu vực WHO Tây Thái Bình Dương.'),
]

d = Document(OUT)

# Find references section start — first ref starts with "American Psychiatric Association"
# Strategy: for each paragraph that matches an EN text, append VN line as new paragraph right after
added = 0
for en, vn in BILING:
    # Find matching paragraph
    for idx, p in enumerate(d.paragraphs):
        if p.text.strip() == en.strip():
            # Insert a new paragraph AFTER this one with Vietnamese text
            new_p = d.paragraphs[idx]._element.addnext(
                d.paragraphs[idx]._element.makeelement(
                    '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}p', {}))
            # Use add_paragraph then move
            new_para = d.add_paragraph()
            r = new_para.add_run(vn)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            # Move the new paragraph's XML element right after the English one
            en_el = p._element
            new_el = new_para._element
            new_el.getparent().remove(new_el)
            en_el.addnext(new_el)
            # Also remove the placeholder we inserted
            try:
                en_el.getnext().getparent().remove(en_el.getnext())
            except Exception:
                pass
            added += 1
            break  # go to next ref

print(f'Bilingual added: {added}/{len(BILING)}')
d.save(OUT)
