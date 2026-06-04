# -*- coding: utf-8 -*-
"""Doc trả lời thầy 2 câu hỏi 28/04/2026:
1. GBD 2019 Mental Disorders Collaborators 2022 + WHO 2022b — đã có bản dịch chưa?
2. β = 0,176 là cao hay trung bình?
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\GBD2019_WHO2022_Beta0176_giai_thich_cho_thay_28042026.docx'
RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def warn(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run('⚠ '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def ok(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run('✓ '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(12); r2.font.name='Times New Roman'

# =====================================================================
title("TRẢ LỜI 2 CÂU HỎI CỦA THẦY", 16)
subtitle("(1) GBD 2019 Mental Disorders Collaborators 2022 + WHO 2022b — đã có bản dịch chưa?")
subtitle("(2) Hệ số β (beta) = 0,176 là cao hay trung bình?")
nr("")
subtitle("Trợ lý nghiên cứu — 28/04/2026 — tiếng Việt thuần — chú thích thuật ngữ tiếng Anh trong ngoặc")

# =====================================================================
# CÂU HỎI 1
# =====================================================================
H1("CÂU 1 — GBD 2019 + WHO 2022b — đã có bản dịch chưa?")

H2("1.1. Trả lời nhanh")
warn("CHƯA — kho 35+ bài hiện có của dự án CHƯA có bản dịch + phản biện đầy đủ cho 2 "
     "tài liệu này. Cả 2 chỉ xuất hiện dưới dạng TRÍCH DẪN trong reference list của các "
     "bài khác đã được dịch.")

H2("1.2. GBD 2019 Mental Disorders Collaborators (2022) là gì?")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài đầy đủ', 'Global, regional, and national burden of 12 mental disorders '
         'in 204 countries and territories, 1990–2019: a systematic analysis for the '
         'Global Burden of Disease Study 2019'],
        ['Tạm dịch tên bài', 'Gánh nặng toàn cầu, khu vực và quốc gia của 12 rối loạn '
         'tâm thần ở 204 quốc gia và vùng lãnh thổ, 1990–2019: phân tích hệ thống cho '
         'Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019'],
        ['Tác giả', 'GBD 2019 Mental Disorders Collaborators (Nhóm nghiên cứu cộng tác '
         'về Rối loạn tâm thần thuộc Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019) — '
         '~150 tác giả từ Viện Sức khỏe Đo lường và Đánh giá (Institute for Health '
         'Metrics and Evaluation — IHME), Đại học Washington, Hoa Kỳ + cộng tác viên '
         'toàn cầu'],
        ['Tạp chí', 'The Lancet Psychiatry'],
        ['Tập / Số / Trang', 'Tập 9(2): trang 137–150'],
        ['Năm + ngày', '2022 — đăng online 11/01/2022'],
        ['DOI', '10.1016/S2215-0366(21)00395-3'],
        ['PMID', '35026139'],
        ['Loại bài', 'Nghiên cứu dịch tễ học toàn cầu — phân tích hệ thống dữ liệu 30 năm'],
        ['Quy mô', 'Đo lường gánh nặng của 12 rối loạn tâm thần ở 204 quốc gia + vùng '
         'lãnh thổ, giai đoạn 1990-2019'],
        ['Trạng thái mở', 'Open Access (CC BY 4.0)'],
        ['Đường dẫn',
         'https://www.thelancet.com/journals/lanpsy/article/PIIS2215-0366(21)00395-3/fulltext'],
    ], [4.0, 12.0])

nr("")
nr("PHÁT HIỆN CHÍNH (tóm lược từ trích dẫn trong các bài khác trong kho dự án):", bold=True)
nr("• Năm 2019: 970 triệu người trên toàn cầu sống với rối loạn tâm thần — tương đương "
   "1/8 dân số thế giới")
nr("• Lo âu + trầm cảm chiếm phần lớn nhất trong gánh nặng (DALY) — tăng ~25% trong "
   "thời kỳ COVID-19 (theo dữ liệu cập nhật)")
nr("• Trẻ em + vị thành niên là nhóm có tỉ lệ rối loạn tâm thần TĂNG NHANH NHẤT")
nr("• Khu vực Mỹ Latinh nhiệt đới + Bắc Mỹ + Châu Đại Dương có gánh nặng cao nhất")
nr("• Bài này là CƠ SỞ DỮ LIỆU TOÀN CẦU CHUẨN — được trích trong hầu hết các bài về "
   "dịch tễ SKTT toàn cầu sau 2022")

H2("1.3. WHO 2022b là gì?")
nr("\"WHO 2022b\" là cách trích dẫn tài liệu thứ 2 của WHO trong cùng năm 2022 — chỉ "
   "có ý nghĩa khi nhìn cùng \"WHO 2022a\" trong cùng bài trích. Trong bối cảnh "
   "Jenkins và cộng sự (2023) — bài đã có trong kho dự án — \"WHO 2022b\" được trích "
   "kèm với GBD 2019 Mental Disorders Collaborators (2022) → KHẢ NĂNG CAO chỉ:",
   italic=True, color=GRAY)
tbl(['Mục', 'Nội dung'],
    [
        ['Tên báo cáo đầy đủ', 'World Mental Health Report: Transforming mental health '
         'for all'],
        ['Tạm dịch', 'Báo cáo Sức khỏe Tâm thần Thế giới: Chuyển đổi sức khỏe tâm thần '
         'cho tất cả mọi người'],
        ['Tổ chức xuất bản', 'World Health Organization (WHO — Tổ chức Y tế Thế giới)'],
        ['Năm + ngày', '2022 — công bố 16/06/2022'],
        ['ISBN', '978-92-4-004933-8 (bản in); 978-92-4-004934-5 (online)'],
        ['Số trang', '~296 trang'],
        ['Loại tài liệu', 'Báo cáo chính thức của WHO (không phải bài báo khoa học) — '
         'tài liệu chính sách + tổng hợp bằng chứng'],
        ['Trạng thái mở', 'Open Access (CC BY-NC-SA 3.0 IGO)'],
        ['Đường dẫn',
         'https://www.who.int/publications/i/item/9789240049338'],
        ['Đặc điểm cấu trúc', '7 chương lớn + nhiều case studies + bằng chứng từ ~200 '
         'quốc gia'],
    ], [4.0, 12.0])

nr("")
nr("PHÁT HIỆN CHÍNH (tóm lược từ trích dẫn):", bold=True)
nr("• Đây là báo cáo SKTT toàn cầu lớn nhất của WHO trong 20 năm (báo cáo trước đó là "
   "World Health Report 2001)")
nr("• Ước tính 25% TĂNG trong tỉ lệ lo âu + trầm cảm toàn cầu trong năm đầu COVID-19 "
   "— phụ nữ + thanh niên là nhóm bị ảnh hưởng nặng nhất (đây chính là số liệu được "
   "trích trong Jenkins 2023)")
nr("• Khoảng cách điều trị (treatment gap): 71% người có rối loạn tâm thần nặng ở các "
   "nước thu nhập thấp — trung bình KHÔNG nhận được điều trị")
nr("• Khuyến nghị 3 con đường thay đổi: (1) Đầu tư vào dịch vụ SKTT cộng đồng; "
   "(2) Bảo vệ + thúc đẩy SKTT; (3) Tăng quyền tự chủ cho người sử dụng dịch vụ")

H2("1.4. Lưu ý về \"WHO 2022a\" và \"WHO 2022b\"")
nr("Cách trích dẫn \"a\", \"b\" được dùng khi cùng một tác giả/tổ chức có nhiều tài "
   "liệu trong cùng năm. Trong Jenkins 2023:")
nr("• WHO 2022a = báo cáo về \"14% thanh thiếu niên (10-19 tuổi) sống với rối loạn "
   "tâm thần năm 2019\" — có thể là Fact Sheet về Adolescent Mental Health trên trang "
   "WHO (cập nhật 2022)")
nr("• WHO 2022b = bài/báo cáo khác — KHẢ NĂNG CAO là World Mental Health Report như "
   "đã mô tả ở mục 1.3")
nr("Để CHẮC CHẮN, em đề xuất kiểm tra reference list đầy đủ của bài Jenkins 2023 (có "
   "trong kho 03_Ban-dich/Dich-11-bai-dau-tien/01_Jenkins_et_al_2023.md) — sẽ thấy "
   "hai entry WHO khác nhau.")

H2("1.5. Đề xuất tiếp theo — nếu thầy muốn em dịch + phản biện đầy đủ")
nr("Em đề xuất 2 phương án — thầy chọn:")
ok("PHƯƠNG ÁN A — Dịch + phản biện đầy đủ cả 2 tài liệu theo workflow chuẩn (giống "
   "đã làm cho 3 bài Brown + Herres 2015 + Steinhoff 2023):")
nr("• Bước 1: Em tải PDF GBD 2019 (open access từ Lancet Psychiatry — 14 trang)")
nr("• Bước 2: Em tải PDF World Mental Health Report 2022 (open access từ WHO — 296 trang)")
nr("• Bước 3: Soạn 2 doc theo cấu trúc 7 phần (trang bìa + thông tin thư mục + bảng "
   "từ viết tắt + dịch song ngữ + bảng dữ liệu Word thật + reference list + phản biện "
   "+ gap)")
nr("• Bước 4: Index vào RAG + KG")
nr("• Bước 5: Kiểm 5-10 vòng")
nr("• Quy mô dự kiến: GBD 2019 ~80-100 KB doc; WHO 2022 ~120-150 KB doc (do nội dung "
   "WHO Report dài hơn)")
nr("")
ok("PHƯƠNG ÁN B — Doc tóm tắt nhanh (~30-40 KB mỗi doc) — chỉ phần chính + reference + "
   "ý nghĩa cho Việt Nam — không dịch song ngữ chi tiết. Phù hợp nếu thầy chỉ cần "
   "biết nội dung tổng quát.")
nr("")
nr("Bác chọn phương án nào? Hoặc bác có muốn em tải PDF trước (vào "
   "02_Papers-goc/) rồi mới quyết định sau?", italic=True, color=GREEN)

# =====================================================================
# CÂU HỎI 2
# =====================================================================
H1("CÂU 2 — Hệ số β (beta) = 0,176 là cao hay trung bình?")

H2("2.1. Trả lời nhanh")
nr("Theo chuẩn DIỄN GIẢI KÍCH THƯỚC HIỆU ỨNG CỦA COHEN (1988) — chuẩn được dùng PHỔ "
   "BIẾN NHẤT trong tâm lý học + dịch tễ học hành vi:", bold=True)
nr("β = 0,176 thuộc khoảng \"NHỎ\" (small effect) — KHÔNG phải trung bình.",
   bold=True, color=BLUE, size=14)
nr("")
nr("Cụ thể:")
tbl(['Khoảng |β| (giá trị tuyệt đối)', 'Mức kích thước hiệu ứng', 'Diễn giải lâm sàng'],
    [
        ['|β| < 0,10', 'KHÔNG đáng kể (negligible)',
         'Có ý nghĩa thống kê nếu n lớn nhưng MAGNITUDE thực tế nhỏ — cần thận trọng'],
        ['0,10 ≤ |β| < 0,30', 'NHỎ (small)',
         'Có ý nghĩa nhưng tác động khiêm tốn — cần kết hợp nhiều yếu tố để có effect rõ'],
        ['0,30 ≤ |β| < 0,50', 'TRUNG BÌNH (medium)',
         'Tác động đáng kể có thể nhận diện được trong thực tế'],
        ['|β| ≥ 0,50', 'LỚN (large)',
         'Tác động mạnh — yếu tố chính giải thích kết quả'],
    ], [4.5, 4.5, 7.0])
nr("")
nr("→ β = 0,176 nằm trong khoảng 0,10–0,30 → mức NHỎ.", bold=True)

H2("2.2. Tại sao phân biệt nhỏ vs trung bình quan trọng?")
nr("Giáo sư Cohen 1988 đặt ngưỡng để giúp nhà nghiên cứu DIỄN GIẢI ý nghĩa thực tế của "
   "kết quả thống kê:")
nr("(1) GIÁ TRỊ p NÓI VỀ Ý NGHĨA THỐNG KÊ; β NÓI VỀ Ý NGHĨA THỰC TẾ. Hai khía cạnh "
   "khác nhau: với cỡ mẫu lớn, ngay cả β = 0,01 cũng có thể có p < 0,05.")
nr("(2) β = 0,176 nghĩa là khi yếu tố tiên lượng X tăng 1 độ lệch chuẩn (standard "
   "deviation — SD), kết quả Y thay đổi ~0,176 độ lệch chuẩn — ~17,6% SD. Trong thực "
   "tế, đây là tác động NHỎ.")
nr("(3) Nếu một bài báo nói \"yếu tố A có tác động đáng kể với β = 0,17\" — nên hiểu "
   "là \"có ý nghĩa thống kê + tác động nhỏ\" chứ KHÔNG phải \"tác động lớn\".")

H2("2.3. Quan điểm cập nhật từ Funder & Ozer (2019) — diễn giải phục hồi cho hiệu ứng nhỏ")
nr("Funder & Ozer (2019) đã CẬP NHẬT diễn giải kích thước hiệu ứng:")
tbl(['|β| hoặc r', 'Funder & Ozer 2019', 'Lý do thay đổi'],
    [
        ['~0,05', 'RẤT NHỎ — chỉ ý nghĩa nếu lặp lại nhiều lần',
         'Trong tâm lý học xã hội, rất khó đo'],
        ['~0,10', 'NHỎ — vẫn có thể có ý nghĩa cho kết quả CỘNG DỒN qua thời gian',
         'Nhiều hiệu ứng nhỏ tích tụ tạo thay đổi lớn'],
        ['~0,20', 'TRUNG BÌNH — đáng chú ý trong nghiên cứu hành vi',
         'Phù hợp với độ phức tạp của hành vi con người'],
        ['~0,30', 'TƯƠNG ĐỐI LỚN — hiếm gặp trong nghiên cứu tâm lý',
         '—'],
        ['≥ 0,40', 'RẤT LỚN — gần như không thực tế trong tâm lý xã hội',
         'Bị nghi ngờ là sai số hoặc thiên lệch'],
    ], [3.0, 5.5, 7.0])
nr("")
nr("→ Theo Funder & Ozer 2019: β = 0,176 thuộc khoảng NHỎ-TRUNG BÌNH (gần ngưỡng giữa "
   "0,10 và 0,20).", bold=True)
nr("Funder & Ozer lập luận rằng tâm lý học XÃ HỘI thường ÁP DỤNG QUÁ NGHIÊM NGẶT chuẩn "
   "Cohen 1988 — nhiều hiệu ứng \"nhỏ\" theo Cohen có thể có ý nghĩa thực tiễn nếu "
   "được tích luỹ qua thời gian + qua nhiều người.")

H2("2.4. Diễn giải thực tế cho β = 0,176 — phụ thuộc bối cảnh")
nr("Cách diễn giải tốt nhất phụ thuộc vào BỐI CẢNH nghiên cứu:")
tbl(['Bối cảnh', 'Diễn giải β = 0,176', 'Ví dụ cụ thể'],
    [
        ['Thử nghiệm có đối chứng ngẫu nhiên (RCT) ngắn hạn',
         'NHỎ — thường cần effect lớn hơn để biện minh chi phí can thiệp',
         'BESST trial 2024 với Cohen d = −0,17 cũng được mô tả là "modest"'],
        ['Nghiên cứu cohort dài hạn (longitudinal cohort study)',
         'NHỎ NHƯNG ĐÁNG CHÚ Ý — tác động qua nhiều năm có thể tích luỹ thành '
         'thay đổi đáng kể',
         'Steinhoff 2023 với β = 0,11–0,15 sau 9 năm — vẫn được tác giả coi là '
         'có ý nghĩa'],
        ['Phân tích gộp (meta-analysis)',
         'NHỎ — nhưng nếu lặp lại nhất quán qua nhiều nghiên cứu, vẫn được coi là '
         'bằng chứng',
         'Phân tích gộp mindfulness 2025 với SMD = −0,14 — vẫn được công bố trên tạp '
         'chí có hệ số tác động cao'],
        ['Y tế công cộng (public health intervention)',
         'NHỎ NHƯNG QUAN TRỌNG — nếu áp dụng cho hàng triệu người, ngay cả tác động '
         'nhỏ cũng tạo ảnh hưởng lớn',
         'Khuyến cáo USPSTF 2022 dựa trên các effects nhỏ-trung bình'],
        ['Nghiên cứu cơ chế sinh học (basic science)',
         'NHỎ — thường kỳ vọng effect lớn hơn (sinh học có cơ chế trực tiếp)',
         '—'],
    ], [4.0, 6.0, 6.0])

H2("2.5. Tóm gọn — em sẽ nói thầy đọc β = 0,176 thế nào?")
ok("Cách diễn giải KHÁCH QUAN nhất cho thầy:")
nr("\"Hệ số β = 0,176 là kích thước hiệu ứng NHỎ theo chuẩn Cohen 1988 — nằm trong "
   "khoảng 0,10–0,30. Tác động này có thể có Ý NGHĨA THỐNG KÊ (nếu p < 0,05) nhưng "
   "MAGNITUDE thực tế khiêm tốn. Diễn giải còn phụ thuộc bối cảnh: trong nghiên cứu "
   "dài hạn hoặc y tế công cộng quy mô lớn, hiệu ứng nhỏ vẫn có thể có ý nghĩa thực "
   "tiễn nếu tích luỹ qua thời gian + qua nhiều người (theo Funder & Ozer 2019). Khi "
   "trích dẫn, nên ghi cả 3 thông tin: β + giá trị p + KTC 95% — tránh chỉ ghi β đơn "
   "lẻ.\"", bold=True)

# =====================================================================
H1("THAM KHẢO ĐẦY ĐỦ")

H2("Cho câu 1 (GBD 2019 + WHO 2022b)")
nr("1. GBD 2019 Mental Disorders Collaborators (2022). Global, regional, and national "
   "burden of 12 mental disorders in 204 countries and territories, 1990-2019: a "
   "systematic analysis for the Global Burden of Disease Study 2019. The Lancet "
   "Psychiatry, 9(2):137-150. DOI: 10.1016/S2215-0366(21)00395-3 — PMID 35026139", size=11)
nr("2. World Health Organization (WHO) (2022). World Mental Health Report: "
   "Transforming mental health for all. Geneva: WHO. ISBN 978-92-4-004933-8. "
   "https://www.who.int/publications/i/item/9789240049338", size=11)
nr("3. World Health Organization (WHO) (2022, cập nhật). Adolescent Mental Health "
   "Fact Sheet. https://www.who.int/news-room/fact-sheets/detail/adolescent-mental-"
   "health (có thể là \"WHO 2022a\")", size=11)

H2("Cho câu 2 (β = 0,176)")
nr("4. Cohen J (1988). Statistical Power Analysis for the Behavioral Sciences. "
   "Lawrence Erlbaum Associates, Hillsdale NJ. ISBN 978-0-8058-0283-2. (Sách kinh điển "
   "đặt nền tảng cho chuẩn diễn giải effect size)", size=11)
nr("5. Funder DC, Ozer DJ (2019). Evaluating effect size in psychological research: "
   "Sense and nonsense. Advances in Methods and Practices in Psychological Science, "
   "2(2):156–168. DOI: 10.1177/2515245919847202 (Cập nhật chuẩn cho tâm lý xã hội)", size=11)
nr("6. Götz FM, Gosling SD, Rentfrow PJ (2022). Small Effects: The Indispensable "
   "Foundation for a Cumulative Psychological Science. Perspectives on Psychological "
   "Science, 17(1):205–215. DOI: 10.1177/1745691620984483 (Bảo vệ giá trị của effect "
   "nhỏ trong khoa học tích luỹ)", size=11)
nr("7. Acock AC (2014). A Gentle Introduction to Stata, 4th ed. Stata Press. (Hướng "
   "dẫn diễn giải β trong hồi quy)", size=11)

H2("Bài liên quan trong kho dự án (đã dịch + phản biện)")
nr("• Brown JSL et al. (2024). BESST trial — Lancet Psychiatry 11(7):504-515. "
   "PMID 38759665 — corpus QT042_BESST. Cohen d = −0,17 cho mẫu chung (small).", size=11)
nr("• Steinhoff A et al. (2023). Longitudinal cohort 9 năm — J Early Adolescence "
   "44(9):1250-1280. PMID 39372429 — β = 0,11-0,15 (small) NHƯNG có ý nghĩa qua "
   "thời gian dài.", size=11)
nr("• Herres J & Ohannessian CM (2015). LPA coping profiles — JAD 186:312-319. "
   "PMID 26275359 — adjusted R² = 0,006-0,074 (rất nhỏ) cho ảnh hưởng của hồ sơ ứng "
   "phó lên triệu chứng.", size=11)

H2("Tài liệu trong kho dự án có TRÍCH DẪN GBD 2019 + WHO 2022")
nr("• Jenkins et al. (2023) — đã dịch ở 03_Ban-dich/Dich-11-bai-dau-tien/"
   "01_Jenkins_et_al_2023.md (trích cả GBD 2019 + WHO 2022a + WHO 2022b)", size=11)
nr("• V-NAMHS 2022 đầy đủ — đã dịch ở 03_Ban-dich/DICH_VNAMHS_2022_DayDu.md (trích "
   "GBD 2019 + COVID-19 Mental Disorders Collaborators 2021 trong phần thảo luận)", size=11)
nr("• GBD ASEAN 2025 (Lancet) — đã dịch ở 03_Ban-dich/DICH_GBD_2021_ASEAN_Lancet_2025."
   "md (trích GBD 2019 trong reference list)", size=11)
nr("• Dong A5 (đã dịch trong dich_A5_Dong.py): \"GBD 2019: trầm cảm chiếm phần lớn "
   "nhất 37,3%, tiếp theo là lo âu 22,9%\"", size=11)

d.save(OUT)
print('Wrote:', OUT)
