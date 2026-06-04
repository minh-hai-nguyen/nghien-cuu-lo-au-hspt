# -*- coding: utf-8 -*-
"""Bài 7: Chen et al. 2023 — theo phong cách CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 7 =====
bl('Tóm tắt bài 7', bold=True)

# BƯỚC 1: ĐỊNH DANH
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình nghiên cứu \u00ab Tỷ lệ và các yếu tố liên quan của triệu chứng trầm cảm và lo âu ở học sinh trung học Trung Quốc \u00bb (Prevalence and Associated Factors of Depressive and Anxiety Symptoms among Chinese Secondary School Students), do Zhangming Chen, Silan Ren, Ruini He và cộng sự (2023), thuộc Khoa Tâm thần, Bệnh viện Tương Nhã 2, Đại học Trung Nam, Trung Quốc, khảo sát trên mẫu 63.205 học sinh trung học cơ sở và trung học phổ thông (tuổi trung bình 14,3) tại thành phố Tự Cống, tỉnh Tứ Xuyên, miền Tây Trung Quốc, xuất bản trên BMC Psychiatry, 23:580.')

# BƯỚC 2: TỔNG QUAN PP + KT1
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng Bảng câu hỏi Sức khỏe Bệnh nhân 9 mục (PHQ-9), Thang đo Rối loạn Lo âu Tổng quát 7 mục (GAD-7), Thang đo Đa chiều về Nạn nhân Bắt nạt Đồng trang lứa (MPVS), Chỉ số Chất lượng Giấc ngủ Pittsburgh (PSQI) và Thang đo Rối loạn Chơi Game Trực tuyến 9 mục-Rút gọn (IGDS9-SF). Nói cách khác, công trình này sử dụng phương pháp sàng lọc cắt ngang quy mô lớn với hồi quy logistic đa biến và điều chỉnh hiệu ứng cụm.', bold=True)

# BƯỚC 3: ĐỊNH NGHĨA (KT5)
bl('PHQ-9 là thang đo 9 mục dựa trên tiêu chí DSM-IV, điểm 0\u201327, ngưỡng ≥10 cho dương tính trầm cảm (Kroenke và cộng sự, 2001). GAD-7 là thang đo 7 mục, điểm 0\u201321, ngưỡng ≥10 cho dương tính lo âu (Spitzer và cộng sự, 2006). Cả hai phiên bản tiếng Trung đã được xác thực cho thanh thiếu niên Trung Quốc.')
bl('MPVS (Multidimensional Peer-Victimization Scale) gồm 16 mục đo 4 loại bắt nạt: thể chất, lời nói, thao túng xã hội và tấn công tài sản. PSQI đo chất lượng giấc ngủ, điểm >5 là rối loạn giấc ngủ. IGDS9-SF đo rối loạn chơi game, điểm ≥32 là dương tính.')

# BƯỚC 4: BIỆN MINH (KT4)
bl('Tổng quan tài liệu của chúng tôi cho thấy nghiên cứu tập trung vào miền Tây Trung Quốc còn rất ít — chỉ 9/51 công bố trong phân tích tổng hợp về trầm cảm thanh thiếu niên Trung Quốc liên quan đến miền Tây. Nghiên cứu trước đây cũng chưa đánh giá toàn diện các yếu tố bắt nạt, giấc ngủ và chơi game trực tuyến cùng lúc.', bold=True)

# BƯỚC 5: LIỆT KÊ DỮ LIỆU
bl('Dữ liệu bao gồm giới tính, tuổi, giai đoạn học tập (THCS/THPT), nơi cư trú (nông thôn/thành thị), loại hình gia đình (hạt nhân/không hạt nhân), con một, trẻ bị bỏ lại (left-behind), bắt nạt đồng trang lứa (4 loại), chất lượng giấc ngủ (PSQI) và rối loạn chơi game trực tuyến (IGDS9-SF).')

# BƯỚC 6: MÔ TẢ QUY TRÌNH
bl('Khảo sát cắt ngang từ tháng 9 đến 12/2020 tại Tự Cống (giai đoạn thuyên giảm COVID-19). Lấy mẫu cụm: chọn 2 quận và 1 huyện, bao gồm tất cả trường trung học trong khu vực. 63.487 học sinh hoàn thành bảng câu hỏi điện tử tại phòng máy tính trường, loại 282 do thiếu dữ liệu (tỷ lệ phản hồi 99,6%). Phân tích bằng R 3.3.3 với gói "survey" để điều chỉnh hiệu ứng cụm và tính tỷ lệ có trọng số. P < 0,001 do cỡ mẫu rất lớn.')

# BƯỚC 9: ĐỊNH LƯỢNG — BẢNG
rb('Kết quả nghiên cứu định lượng')

bl('Tỷ lệ có trọng số của triệu chứng trầm cảm và lo âu (N = 63.205):', bold=True)
tbl(
    ['Tình trạng', 'Tỷ lệ có trọng số', 'KTC 95%'],
    [
        ['Trầm cảm (PHQ-9 ≥10)', '23,0%', '19,6\u201327,0%'],
        ['Lo âu (GAD-7 ≥10)', '13,9%', '11,2\u201317,0%'],
        ['Đồng mắc cả hai', '11,5%', '9,0\u201315,0%'],
        ['Trầm cảm hoặc lo âu', '25,5%', '21,8\u201330,0%'],
    ],
    widths=[5.0, 3.5, 3.5]
)
doc.add_paragraph()

bl('Yếu tố liên quan độc lập (hồi quy logistic đa biến):', bold=True)
tbl(
    ['Yếu tố', 'OR đa biến', 'KTC 95%', 'Mức liên quan'],
    [
        ['Rối loạn giấc ngủ (PSQI >5)', '6,99', '6,69\u20137,30', 'Mạnh'],
        ['Rối loạn chơi game (IGD)', '5,00', '4,42\u20135,66', 'Mạnh'],
        ['Thao túng xã hội', '1,97', '1,87\u20132,08', 'Trung bình'],
        ['Bắt nạt lời nói', '1,70', '1,60\u20131,80', 'Trung bình'],
        ['Giới tính nữ', '1,55', '1,49\u20131,62', 'Trung bình'],
        ['Bắt nạt thể chất', '1,51', '1,44\u20131,59', 'Trung bình'],
        ['Gia đình không hạt nhân', '1,31', '1,25\u20131,38', 'Yếu'],
        ['Tấn công tài sản', '1,25', '1,18\u20131,33', 'Yếu'],
        ['Sống thành thị', '1,11', '1,06\u20131,17', 'Yếu'],
        ['Con một', '1,06', '1,01\u20131,12', 'Yếu'],
    ],
    widths=[4.5, 2.5, 3.0, 2.5]
)
doc.add_paragraph()

# BƯỚC 10: ĐỐI CHIẾU (KT3)
bh3('Khác biệt giới tính : Nữ giới luôn được báo cáo có nguy cơ trầm cảm và lo âu cao hơn nam giới (Salk và cộng sự, 2017). Thực tế, phân tích tổng hợp thanh thiếu niên Trung Quốc xác nhận xu hướng này. Chúng tôi xác nhận trong mẫu 63.205 học sinh rằng nữ chiếm 58% nhóm căng thẳng so với 48% nhóm không căng thẳng (OR = 1,55, P < 0,001). Đáng chú ý, rối loạn giấc ngủ (OR = 6,99) và rối loạn chơi game (OR = 5,00) là hai yếu tố mạnh nhất \u2014 vượt trội hơn mọi yếu tố nhân khẩu học.')

# BƯỚC 11+12: NHẬN XÉT
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về vai trò trung tâm của giấc ngủ, game và bắt nạt')
bl('*Rối loạn giấc ngủ \u2014 yếu tố nguy cơ mạnh nhất.* Nổi bật nhất là rối loạn giấc ngủ (PSQI >5) có OR = 6,99 \u2014 yếu tố nguy cơ mạnh nhất trong tất cả các biến. 66% học sinh có căng thẳng tâm thần có rối loạn giấc ngủ, so với 17% nhóm không căng thẳng. Phát hiện này tương tự Zhu và cộng sự (2025): ngủ <5 giờ có AOR = 13,71 cho trầm cảm tại Trung Quốc. Giấc ngủ là yếu tố bảo vệ có thể can thiệp được.')
bl('*Rối loạn chơi game trực tuyến.* IGD (IGDS9-SF ≥32) có OR = 5,00 \u2014 yếu tố nguy cơ mạnh thứ hai. 8,3% học sinh căng thẳng có triệu chứng IGD so với 1,0% nhóm không căng thẳng. Đây là bằng chứng mạnh cho mối quan hệ giữa chơi game gây nghiện và sức khỏe tâm thần, tương tự phát hiện của Saikia và cộng sự (2023) về game tại Ấn Độ.')
bl('*Bắt nạt đồng trang lứa \u2014 4 loại.* Tất cả 4 loại bắt nạt đều liên quan có ý nghĩa, nhưng thao túng xã hội (OR = 1,97) và bắt nạt lời nói (OR = 1,70) mạnh hơn bắt nạt thể chất (OR = 1,51) và tấn công tài sản (OR = 1,25). Phân tích chi tiết theo loại bắt nạt là điểm mạnh phương pháp luận, ít nghiên cứu nào thực hiện.')

# BƯỚC 13: KẾT LUẬN (Kiến trúc B)
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy tỷ lệ trầm cảm 23,0% và lo âu 13,9% ở 63.205 học sinh trung học miền Tây Trung Quốc \u2014 cỡ mẫu lớn nhất trong số 11 bài nghiên cứu, gợi ý rằng giấc ngủ và chơi game là hai mục tiêu can thiệp ưu tiên (OR = 6,99 và 5,00). Bắt nạt học đường, đặc biệt thao túng xã hội và bắt nạt lời nói, cũng đóng vai trò quan trọng. Trong bối cảnh miền Tây Trung Quốc \u2014 nơi điều kiện kinh tế kém hơn miền Đông, chiến lược can thiệp cần tập trung vào cải thiện giấc ngủ, kiểm soát chơi game, và xây dựng môi trường trường học an toàn.', bold=True)

# PHẢN BIỆN
rh2('Phản biện')
bl('Điểm mạnh nổi bật: cỡ mẫu cực lớn (63.205), tỷ lệ phản hồi 99,6%, BMC Psychiatry Q1 (PubMed/Scopus-indexed), sử dụng 5 công cụ đã xác thực, hồi quy logistic đa biến với điều chỉnh hiệu ứng cụm, phân tích phân nhóm THCS/THPT. Tuy nhiên, thiết kế cắt ngang không cho phép suy luận nhân quả \u2014 rối loạn giấc ngủ có thể là hậu quả chứ không chỉ là nguyên nhân của trầm cảm/lo âu. Dữ liệu thu thập trong giai đoạn thuyên giảm COVID-19 (T9-12/2020) có thể bị ảnh hưởng bởi hiệu ứng hậu đại dịch. Sử dụng ngưỡng cắt ≥10 cho PHQ-9 và GAD-7 (trung bình-nặng) cho kết quả bảo thủ hơn so với ngưỡng ≥5 \u2014 điều này giải thích tỷ lệ thấp hơn so với các nghiên cứu khác.')

# HƯỚNG NC
rh2('Hướng nghiên cứu tiếp theo')
bl('Nghiên cứu dọc để xác định chiều hướng nhân quả giữa rối loạn giấc ngủ, chơi game và trầm cảm/lo âu. Nghiên cứu can thiệp cải thiện giấc ngủ (vệ sinh giấc ngủ) và kiểm soát thời gian chơi game tại trường. Mở rộng đến các vùng khác của miền Tây Trung Quốc để đánh giá tính khái quát. So sánh kết quả với miền Đông Trung Quốc (kinh tế phát triển hơn) để đánh giá ảnh hưởng kinh tế xã hội.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. Cỡ mẫu 63.205, BMC Psychiatry Q1, 5 công cụ xác thực, hồi quy đa biến có điều chỉnh cụm, phân tích phân nhóm.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('07_Chen_et_al_2023.docx')
print('07_Chen_et_al_2023.docx OK')
