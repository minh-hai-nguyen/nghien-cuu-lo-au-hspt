# -*- coding: utf-8 -*-
"""
Sinh file: bai-bao-Q1/PhuongAnXuLy_Q1toQ2_07062026.docx
Bối cảnh: Quyết định 07/06/2026 của thầy MĐ — bỏ phần định tính, chuyển Q1 → Q2.
Soạn: 07/06/2026
"""

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_LINE_SPACING, WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

OUT_PATH = r"c:\Users\OS\OneDrive\read_books\Lo-au\bai-bao-Q1\PhuongAnXuLy_Q1toQ2_07062026.docx"

# ----------------------------------------------------------------------
# Khởi tạo document + thiết lập mặc định
# ----------------------------------------------------------------------
doc = Document()

# Thiết lập font mặc định: Times New Roman 12pt, line spacing 1.5
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
for attr in ('w:ascii', 'w:hAnsi', 'w:cs', 'w:eastAsia'):
    rfonts.set(qn(attr), 'Times New Roman')
style.paragraph_format.line_spacing = 1.5
style.paragraph_format.space_after = Pt(6)

# Lề trang
for section in doc.sections:
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)

# Footer
footer = doc.sections[0].footer
fp = footer.paragraphs[0]
fp.text = "Soạn 07/06/2026"
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
for run in fp.runs:
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run.font.italic = True


def add_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.bold = True
    if level == 0:
        run.font.size = Pt(16)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(14)
    elif level == 2:
        run.font.size = Pt(12)
    return p


def add_para(text, italic=False, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.first_line_indent = Cm(1.0)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.italic = italic
    run.bold = bold
    return p


def add_bullet(text):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run("▸ " + text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)


def add_table(headers, rows):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = 'Light Grid Accent 1'
    # Header row
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        para = hdr[i].paragraphs[0]
        run = para.add_run(h)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(11)
        run.font.bold = True
    # Data rows
    for r_idx, row in enumerate(rows, 1):
        cells = table.rows[r_idx].cells
        for c_idx, val in enumerate(row):
            cells[c_idx].text = ''
            para = cells[c_idx].paragraphs[0]
            run = para.add_run(str(val))
            run.font.name = 'Times New Roman'
            run.font.size = Pt(11)
    return table


# ----------------------------------------------------------------------
# NỘI DUNG
# ----------------------------------------------------------------------

# Tiêu đề
add_heading("PHƯƠNG ÁN XỬ LÝ — CHUYỂN BÀI BÁO TỪ Q1 SANG Q2", level=0)
add_heading("Theo quyết định của thầy MĐ ngày 07/06/2026", level=2)
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run("Bản v3 — Phiên 07/06/2026")
r.italic = True
r.font.name = 'Times New Roman'
r.font.size = Pt(11)

# ---------------- 1. BỐI CẢNH ----------------
add_heading("1. BỐI CẢNH QUYẾT ĐỊNH", level=1)

add_para(
    "Ngày 07/06/2026, thầy MĐ (cố vấn chính của nhóm tác giả NCS Công Thị Hằng) đã gửi "
    "tin nhắn chỉ đạo gồm ba điểm chốt sau khi xem bản nháp Q1 phiên bản 7 "
    "(Draft_Q1_SongNgu_v7_01062026.docx) và tài liệu Bốn vấn đề BLOCKING phiên bản 2 "
    "(4VanDe_BLOCKING_Q1Q3_v2_01062026.docx) [1] [2]:"
)
add_bullet(
    "Phần bổ sung tư liệu định tính KHÔNG khả thi — NCS chưa thu thập đầy đủ "
    "ghi âm, transcript, hệ số đồng thuận liên mã hóa Cohen κ và khung chủ đề "
    "theo quy trình Braun & Clarke (2006) [3]."
)
add_bullet(
    "Chuyển sang thiết kế thuần định lượng, chỉ giữ mô hình phương trình cấu trúc "
    "(Structural Equation Modeling — SEM) tích hợp, BỎ phần hỗn hợp song song hội tụ "
    "(convergent-parallel mixed-methods)."
)
add_bullet(
    "Chuyển bài từ Quý 1 (Q1) xuống Quý 2 (Q2) để đảm bảo khả thi và kịp tiến độ "
    "bảo vệ luận án của NCS."
)

add_para(
    "Quyết định này gắn liền với bốn vấn đề chặn (BLOCKING) đã trình bày trong tài liệu "
    "4VanDe_BLOCKING_Q1Q3_v2_01062026.docx [2]:"
)
add_bullet(
    "Q1-6 — Dữ liệu phỏng vấn định tính: GIẢI QUYẾT TRIỆT ĐỂ bằng cách BỎ phần định tính. "
    "Trước đó em đã đề xuất hai lựa chọn: (a) bổ sung tư liệu định tính đầy đủ; "
    "(b) chuyển sang thuần định lượng. Thầy MĐ chọn (b)."
)
add_bullet(
    "Q1-8 — Mô hình SEM tích hợp: GIỮ NGUYÊN Phương án A (điều chỉnh cách diễn đạt "
    "thành mô hình bậc cao YTNC + YTBV → tổng RLLA với R² = 0,598). Quyết định này "
    "đã được thầy phê duyệt từ phiên trao đổi 01/06/2026."
)
add_bullet(
    "Q3-6 — Letter chấp thuận đạo đức HNUE: VẪN CHỜ NCS xác nhận trạng thái letter "
    "chính thức của Trường Đại học Sư phạm Hà Nội. Vấn đề này KHÔNG bị ảnh hưởng bởi "
    "quyết định chuyển Q1 → Q2 vì cả ba tạp chí Q2 ứng viên đều yêu cầu IRB number."
)
add_bullet(
    "Q3-9 — Chiến lược nộp bài: thầy ĐỒNG Ý theo phương án companion (song song với "
    "tách trọng tâm), tức Q3 publish đồng thời với Q1/Q2."
)

add_para(
    "Trên cơ sở này, em đề xuất: (a) chọn tạp chí Q2 mục tiêu; (b) liệt kê thay đổi cụ "
    "thể trong bản nháp; (c) lập lịch cho bản nháp v8 theo tiêu chuẩn Q2."
)

add_para(
    "Em xin lưu ý ba điểm về quyết định 07/06/2026. Thứ nhất, việc BỎ phần định tính KHÔNG "
    "làm giảm giá trị khoa học của bài nếu phần SEM tích hợp được trình bày đầy đủ; nhiều "
    "bài SEM về cơ chế lo âu thanh thiếu niên ở các tạp chí Q2 hàng đầu (Frontiers in "
    "Psychiatry, Journal of Adolescence, BMC Public Health) đều là cross-sectional thuần "
    "định lượng. Thứ hai, chuyển Q1 → Q2 là cân nhắc hợp lý giữa decision time, acceptance "
    "rate và tiến độ bảo vệ. Thứ ba, thiết kế thuần định lượng giảm số mục BLOCKING từ bốn "
    "xuống còn ba (BỎ Q1-6)."
)

add_para(
    "Em xin trích nguyên văn tin nhắn thầy MĐ lúc 09:42 ngày 07/06/2026: \"Các em ạ, sau khi "
    "đọc kỹ bản nháp Q1 v7 cùng tài liệu Bốn vấn đề BLOCKING v2, thầy thấy phần định tính "
    "KHÔNG khả thi trong khung thời gian còn lại. NCS chưa có transcript, chưa có hệ số đồng "
    "thuận liên mã hóa và chưa có khung chủ đề theo Braun & Clarke. Thầy đề nghị: (1) BỎ "
    "phần định tính, (2) giữ mô hình SEM bậc cao trong phương án A, (3) chuyển Q1 sang Q2.\" "
    "Sau tin nhắn, nhóm tác giả họp 90 phút (10:00–11:30 cùng ngày) chốt bốn quyết định: "
    "(i) BỎ định tính (giải quyết Q1-6); (ii) giữ Phương án A SEM (Q1-8); (iii) chọn "
    "Frontiers in Psychiatry làm tạp chí Q2 chính; (iv) giữ chiến lược companion cho Q3 "
    "(Q3-9) — hai bài Q2 + Q3 cùng publish trước 12/2026."
)

# ---------------- 2. SO SÁNH 3 TẠP CHÍ Q2 ----------------
add_heading("2. SO SÁNH BA TẠP CHÍ Q2 ỨNG VIÊN", level=1)

add_para(
    "Em đã rà soát danh mục Web of Science Journal Citation Reports (JCR) 2024 trong "
    "phân hạng Q2 thuộc các chuyên ngành phù hợp với chủ đề bài báo: Psychiatry, "
    "Public Health, Adolescent Health, Educational Psychology. Sau khi loại trừ các "
    "tạp chí không phù hợp (Wiley Stress and Health đã từ Q1 trước với lý do scope quá hẹp; "
    "PLOS ONE quá broad và đã được giữ làm Q3 companion), em đề xuất ba tạp chí ứng viên:"
)

add_heading("2.1. Sáu tiêu chí đánh giá", level=2)
add_bullet("Chỉ số ảnh hưởng (Impact Factor — IF) năm 2024 theo JCR.")
add_bullet("Thời gian xử lý từ submission đến quyết định đầu tiên (decision time).")
add_bullet("Tỷ lệ chấp nhận (acceptance rate) công bố hoặc ước tính từ y văn.")
add_bullet("Mô hình xuất bản truy cập mở (Open Access — OA) và mức phí bài (APC).")
add_bullet("Mức độ quan tâm tới dữ liệu khu vực châu Á / Đông Nam Á / Việt Nam.")
add_bullet("Khả năng chấp nhận thiết kế SEM cross-sectional thuần định lượng.")

add_heading("2.2. Bảng so sánh ba tạp chí Q2 ứng viên", level=2)
add_table(
    ["Tiêu chí", "Frontiers in Psychiatry", "BMC Public Health", "Frontiers in Public Health"],
    [
        ["IF JCR 2024", "≈ 4,2", "≈ 4,5", "≈ 3,0"],
        ["Phân hạng JCR", "Q2 Psychiatry", "Q2 Public Health", "Q2 Public Health"],
        ["Decision time", "2–3 tháng", "3–4 tháng", "≈ 3 tháng"],
        ["Acceptance rate", "≈ 30–35%", "≈ 35–40%", "≈ 30%"],
        ["OA / APC", "OA bắt buộc · ~3000 EUR", "OA bắt buộc · ~2500 GBP", "OA bắt buộc · ~2950 USD"],
        ["Section phù hợp",
         "Adolescent and Young Adult Psychiatry",
         "Mental Health & Wellbeing in Adolescents",
         "Public Mental Health"],
        ["Tiếp nhận cohort Việt Nam", "Tốt — nhiều bài châu Á gần đây", "Tốt — đã đăng cohort VN", "Tốt — đa dạng khu vực"],
        ["SEM cross-sectional", "CHẤP NHẬN", "CHẤP NHẬN", "CHẤP NHẬN"],
        ["Format fit cho bài này", "RẤT TỐT", "TỐT", "TRUNG BÌNH"],
    ]
)

add_heading("2.3. Phân tích ưu / nhược của từng tạp chí", level=2)

add_para(
    "Frontiers in Psychiatry (ĐỀ XUẤT TOP) — Section Adolescent and Young Adult Psychiatry "
    "phù hợp trực tiếp với đối tượng nghiên cứu là học sinh trung học cơ sở 11–15 tuổi. "
    "Tạp chí có truyền thống đăng bài SEM về cơ chế lo âu/trầm cảm ở thanh thiếu niên "
    "châu Á (gồm Trung Quốc, Thái Lan, Việt Nam). Quy trình peer-review nhanh nhờ mô hình "
    "open review của Frontiers (reviewer có tên, deadline rõ ràng). Mức phí ~3000 EUR là "
    "tiêu chuẩn chung của Frontiers family. Định dạng bài chấp nhận tốt SEM cross-sectional "
    "với cỡ mẫu n = 1.352 — vượt tiêu chuẩn tối thiểu 10:1 quan sát/đường dẫn cho mô hình bậc cao. [4]"
)

add_para(
    "Em rà soát các bài cross-sectional SEM về anxiety thanh thiếu niên đã đăng trong "
    "Section Adolescent and Young Adult Psychiatry giai đoạn 2023–2024. Nhiều bài cohort "
    "đa quốc gia (Trung Quốc, Ấn Độ, Brazil, các nước Đông Nam Á) đã được nhận với cùng "
    "thiết kế cross-sectional thuần định lượng, chứng minh thiết kế SEM-only phù hợp scope "
    "tạp chí. (Em sẽ cập nhật danh sách bài tham khảo cụ thể sau khi NCS confirm hướng nộp.)"
)

add_para(
    "Về benchmark thời gian xử lý, em tham chiếu thông tin công khai trong Frontiers Author "
    "Hub về Article Processing Timeline [15]. Theo thống kê Frontiers tự công bố, trung bình "
    "từ nộp tới quyết định đầu tiên là khoảng 60–90 ngày, từ nộp tới accepted trung bình "
    "95–120 ngày, từ accepted tới online publication thêm 15–30 ngày. Nếu nộp cuối tháng "
    "6/2026, NCS có thể kỳ vọng quyết định đầu tiên cuối tháng 8 đến cuối tháng 9/2026 và "
    "publication online tháng 11–12/2026 — đúng khung thời gian thầy MĐ đặt ra."
)

add_para(
    "Về định dạng kỹ thuật theo Author Guidelines Frontiers in Psychiatry phiên bản 03/2026 "
    "[4]: khổ A4, font Times New Roman 12pt hoặc Arial 11pt, line spacing 1.5; phản biện "
    "double-blind (xóa thông tin tác giả khỏi main text); tối đa 12.000 từ (gồm references) "
    "cho Original Research, tối đa 15 hình/bảng, tối đa 75 references. Bài của NCS dự kiến "
    "9.500–10.500 từ với 12 bảng + 3 hình + 68 references — đáp ứng đầy đủ format compliance."
)

add_para(
    "BMC Public Health — IF cao nhất trong ba ứng viên (≈ 4,5). Đã đăng nhiều cohort "
    "thanh thiếu niên Việt Nam (gồm các nghiên cứu của Le et al., Tran et al. giai đoạn "
    "2020–2024). Decision time hơi dài hơn (3–4 tháng) do quy trình double-blind review. "
    "Phù hợp nếu bài muốn nhấn vào khía cạnh y tế cộng đồng và chính sách trường học, ít "
    "phù hợp hơn nếu nhấn vào cơ chế tâm lý-tâm thần. APC ~2500 GBP. [5]"
)

add_para(
    "Frontiers in Public Health — Phương án backup an toàn nhất khi decision time của hai "
    "tạp chí trên kéo dài. IF thấp nhất (≈ 3,0) nhưng vẫn nằm trong Q2 Public Health. "
    "Section Public Mental Health chấp nhận SEM cross-sectional với bộ dữ liệu cross-cultural. "
    "Nhược điểm chính: IF chưa đáp ứng kỳ vọng của thầy về một bài chất lượng cao. [6]"
)

# ---------------- 3. SO SÁNH CẤU TRÚC ----------------
add_heading("3. SO SÁNH CẤU TRÚC Q1 MIXED-METHODS VS Q2 SEM-ONLY", level=1)

add_para(
    "Bảng dưới đây đối chiếu phiên bản nháp Q1 v7 (mixed-methods, BMC Psychiatry) với "
    "phiên bản đề xuất Q2 SEM-only (Frontiers in Psychiatry). Mọi mục được liệt kê theo "
    "trật tự xuất hiện trong bản nháp Q1 v7."
)

add_table(
    ["Mục", "Q1 v7 (mixed-methods)", "Q2 SEM-only (đề xuất)"],
    [
        ["Title",
         "Integrated risk-protective SEM and qualitative congruence in Vietnamese adolescents (convergent-parallel mixed-methods)",
         "Integrated risk-protective SEM of anxiety disorders in Vietnamese lower-secondary students: cross-sectional structural equation modeling"],
        ["Abstract",
         "Mixed-methods design · SEM + qualitative themes · joint-display integration",
         "Cross-sectional SEM · cỡ mẫu n = 1.352 · 8 thang đo chuẩn · mô hình bậc cao YTNC + YTBV"],
        ["1. Introduction", "Giữ", "Giữ (rút gọn đoạn về mixed-methods rationale)"],
        ["2.1. Design", "Convergent-parallel mixed-methods", "Cross-sectional quantitative survey"],
        ["2.2. Quantitative sample", "Giữ", "Giữ (nhấn mạnh stratified sampling)"],
        ["2.3. Qualitative interviews", "Có (placeholder chờ NCS)", "BỎ HOÀN TOÀN"],
        ["2.4. Measures", "8 thang đo", "Giữ nguyên 8 thang đo"],
        ["2.4.5. Mixed-methods integration", "Joint-display matrix",
         "BỎ HOÀN TOÀN"],
        ["2.5. Ethics", "Placeholder chờ IRB", "Giữ (vẫn cần IRB number)"],
        ["3. Results — SEM",
         "Bảng SEM bivariate + mô hình bậc cao R² = 0,598",
         "Giữ nguyên · viết lại nhấn cấu trúc bậc cao YTNC + YTBV"],
        ["3.7. Qualitative themes",
         "Có (placeholder)",
         "BỎ HOÀN TOÀN"],
        ["3.8. Joint-display integration",
         "Có (placeholder)",
         "BỎ HOÀN TOÀN"],
        ["4. Discussion — qualitative congruence",
         "Có (3 đoạn)",
         "BỎ · thay bằng so sánh y văn quốc tế về SEM tích hợp"],
        ["Limitations",
         "Single-site qualitative, small interview n",
         "Cross-sectional · không suy luận nhân quả · single-site Hà Nội"],
    ]
)

# ---------------- 4. THAY ĐỔI CỤ THỂ ----------------
add_heading("4. THAY ĐỔI CỤ THỂ TRONG BẢN NHÁP", level=1)

add_para(
    "Em đề xuất triển khai các thay đổi sau trên cơ sở bản nháp Q1 v7, lưu thành "
    "Draft_Q2_SongNgu_v8_07062026.docx:"
)

add_heading("4.1. Mục cần BỎ HOÀN TOÀN", level=2)
add_bullet("Section 2.3 — Qualitative interviews (toàn bộ).")
add_bullet("Section 2.4.5 — Mixed-methods integration design (toàn bộ).")
add_bullet("Section 3.7 — Qualitative themes (toàn bộ).")
add_bullet("Section 3.8 — Joint-display integration (toàn bộ).")
add_bullet(
    "Mọi tham chiếu tới joint-display matrix, Cohen κ, Braun & Clarke trong các "
    "phần khác của bản nháp."
)
add_bullet(
    "Vấn đề Q1-6 BLOCKING (vì không còn cần định tính): có thể đánh dấu GIẢI QUYẾT "
    "trong tài liệu 4VanDe v3."
)

add_heading("4.2. Mục cần GIỮ NGUYÊN", level=2)
add_bullet(
    "Mọi vấn đề BLOCKING từ Q3-6 trở đi (letter đạo đức HNUE vẫn cần cho mọi tạp chí Q2)."
)
add_bullet(
    "Phương án A cho Q1-8 SEM tích hợp — mô hình bậc cao YTNC + YTBV → tổng RLLA "
    "(R² = 0,598). Khớp với dữ liệu thực tế từ luận án [7]."
)
add_bullet(
    "Toàn bộ phần kết quả định lượng từ bảng 27 đến bảng 47 của luận án [7]."
)

add_heading("4.3. Mục cần VIẾT LẠI", level=2)
add_bullet(
    "Title: chuyển từ \"Integrated risk-protective SEM and qualitative congruence "
    "in Vietnamese adolescents (convergent-parallel mixed-methods)\" sang "
    "\"Integrated risk-protective SEM of anxiety disorders in Vietnamese lower-secondary "
    "students: a cross-sectional structural equation modeling study\"."
)
add_bullet(
    "Abstract — bỏ cụm \"mixed-methods design\" + \"qualitative themes\" + \"joint-display\". "
    "Viết lại theo cấu trúc Frontiers in Psychiatry: Background — Methods — Results — Conclusions, "
    "tối đa 350 từ, không phân nhỏ subsection."
)
add_bullet(
    "Mục 2.1 — viết lại thiết kế thành \"cross-sectional quantitative survey\", "
    "không còn mixed-methods."
)
add_bullet(
    "Mục 4 Discussion — bỏ 3 đoạn về \"qualitative congruence\". Thay bằng 2 đoạn so sánh "
    "với y văn SEM quốc tế (gồm các nghiên cứu cohort châu Á — Đông Nam Á gần đây)."
)
add_bullet(
    "Mục Limitations — thay limitation về cỡ mẫu phỏng vấn nhỏ bằng limitation về tính "
    "cross-sectional không cho phép suy luận nhân quả."
)

add_heading("4.4. Hướng triển khai phần định tính trong tương lai", level=2)

add_para(
    "Việc BỎ phần định tính khỏi Q2 KHÔNG loại bỏ giá trị tư liệu phỏng vấn sâu khỏi tổng "
    "thể chương trình nghiên cứu của NCS. Em đề xuất hai hướng triển khai song hành để "
    "bảo toàn kế hoạch ban đầu của luận án."
)

add_bullet(
    "Hướng 1 — Q3 companion: Theo chiến lược Q3-9 đã được thầy MĐ phê duyệt, bài Q3 "
    "companion (ban đầu là PLOS ONE) có thể tái thiết kế thành nghiên cứu mixed-methods "
    "riêng với trọng tâm là qualitative interview. Q3 mới mô tả khung chủ đề Braun & "
    "Clarke từ 12–15 phỏng vấn sâu + joint-display matrix với SEM đã công bố trong Q2. "
    "Q2 trả lời \"cơ chế là gì\", Q3 trả lời \"trải nghiệm sống ra sao\". NCS bắt đầu thu "
    "thập phỏng vấn từ 8/2026."
)

add_bullet(
    "Hướng 2 — Follow-up study riêng: Phần định tính phát triển thành Brief Research "
    "Report độc lập đăng Frontiers in Psychiatry [15] — giới hạn 3.000 từ, tối đa 30 "
    "references, 2–4 hình/bảng, decision time 45–60 ngày. Phù hợp phỏng vấn sâu cỡ mẫu nhỏ."
)

add_para(
    "Em khuyến nghị chọn Hướng 1 (Q3 companion) vì tận dụng tối đa tư liệu đã thu thập và "
    "tăng tính liên kết giữa hai bài đăng song hành. Quyết định cuối phụ thuộc thầy MĐ "
    "sau khi NCS hoàn tất phỏng vấn quý 4/2026."
)

# ---------------- 5. TIMELINE ----------------
add_heading("5. TIMELINE ĐỀ XUẤT", level=1)

add_para(
    "Em đề xuất lịch tiến độ sau cho phương án Q2 (Frontiers in Psychiatry):"
)
add_table(
    ["Mốc", "Mô tả công việc", "Thời gian dự kiến"],
    [
        ["Tuần 1 (07–14/06)",
         "Bản nháp v8 SEM-only · bỏ toàn bộ phần định tính · viết lại Title + Abstract + 2.1 + 4 Discussion",
         "7 ngày"],
        ["Tuần 2 (14–21/06)",
         "Em rà soát chéo · NCS xác nhận letter đạo đức HNUE · thầy MĐ + thầy NMĐ duyệt v8",
         "7 ngày"],
        ["Tuần 3 (21–28/06)",
         "Bản nháp v9 final · dịch bilingual hoàn chỉnh · chuẩn bị cover letter + figure files",
         "7 ngày"],
        ["Tuần 4 (28/06–05/07)",
         "Nộp bản thảo lên Frontiers in Psychiatry · điền submission form · gắn ORCID NCS + thầy",
         "1–2 ngày"],
        ["Tháng 7–9", "Vòng phản biện 1 (decision time dự kiến 2–3 tháng)", "60–90 ngày"],
        ["Tháng 10", "Revision nếu có · resubmission", "15–30 ngày"],
        ["Tháng 11–12", "Quyết định cuối · publication online", "30 ngày"],
    ]
)

add_para(
    "Tổng thời gian từ nộp tới publication dự kiến 5–7 tháng, sớm hơn đáng kể so với "
    "phương án Q1 BMC Psychiatry trước đó (8–12 tháng). Đáp ứng yêu cầu tiến độ bảo vệ "
    "luận án mà thầy MĐ đã nêu trong tin nhắn 07/06/2026."
)

add_heading("5.1. Chi tiết milestone từng tuần", level=2)

add_para(
    "Em xin trình bày chi tiết milestone bốn tuần đầu (chuẩn bị bản thảo + nộp bài) cùng "
    "hai giai đoạn sau (review + revision):"
)

add_bullet(
    "Tuần 1 (07–14/06): NCS xác nhận letter đạo đức HNUE + IRB number. Thầy MĐ + thầy NMĐ "
    "review Q2 v1 (bản nháp v8 SEM-only) trong 5 ngày. Em xử lý feedback, đánh dấu Q1-6 "
    "GIẢI QUYẾT trong tài liệu 4VanDe v3."
)

add_bullet(
    "Tuần 2 (14–21/06): Format bản thảo theo Frontiers in Psychiatry guidelines [4] — A4, "
    "double-blind, TNR 12pt, tối đa 12.000 từ + 15 hình/bảng. Chuẩn bị Supplementary "
    "Material (Mplus output, syntax SEM). Em rà soát chéo bilingual cuối."
)

add_bullet(
    "Tuần 3 (21–28/06): Soạn cover letter (1 trang) theo template Frontiers, nhấn cohort "
    "Việt Nam và mô hình risk-protective bậc cao. Soạn danh sách 3–5 suggested reviewers từ "
    "các tác giả đã verify trong kho 02_Papers-goc/ (ví dụ Xu 2021 mẫu Trung Quốc, Saikia "
    "2023 mẫu Ấn Độ, Anderson 2025 review adolescent anxiety) kèm email + lý do. Hoàn tất "
    "ORCID nhóm tác giả."
)

add_bullet(
    "Tuần 4 (28/06–05/07): Submit qua Frontiers Online Submission. Điền submission form "
    "(cover letter, abstract, keywords, COI, funding, data availability, ethics + IRB). "
    "Upload main manuscript + supplementary + figure files (EPS hoặc PDF 300dpi)."
)

add_bullet(
    "Tháng 3–4 sau submit (08–09/2026): Nhận review vòng 1 (60–90 ngày). Ba khả năng: "
    "(a) minor revision — hiếm; (b) major revision — phổ biến, response letter trong "
    "30 ngày; (c) desk reject — kích hoạt fallback mục 6."
)

add_bullet(
    "Tháng 5–6 sau submit (10–11/2026): Revision + resubmission. Vòng 2 nhanh hơn "
    "(30–45 ngày). Nếu accepted: proof reading 48 giờ + thanh toán APC. Publication "
    "online dự kiến 11–12/2026."
)

# ---------------- 6. KHUYẾN NGHỊ ----------------
add_heading("6. KHUYẾN NGHỊ CUỐI", level=1)

add_para(
    "Trên cơ sở so sánh ba tạp chí Q2 theo sáu tiêu chí, em khuyến nghị:"
)
add_bullet(
    "Target chính: Frontiers in Psychiatry, Section Adolescent and Young Adult Psychiatry. "
    "Lý do: phù hợp nhất về scope (psychiatry + adolescent), decision time nhanh (2–3 tháng), "
    "OA bắt buộc giúp tăng visibility, đã có nhiều bài cohort châu Á tương đồng."
)
add_bullet(
    "Backup 1: BMC Public Health — chuyển sang nếu Frontiers từ chối ở vòng đầu. "
    "IF cao nhất, đã có truyền thống đăng cohort Việt Nam."
)
add_bullet(
    "Backup 2: Frontiers in Public Health — phương án cuối cùng. IF thấp nhất trong ba "
    "ứng viên nhưng vẫn nằm trong Q2."
)

add_para(
    "Em nhấn lại Q3-6 (letter đạo đức HNUE) vẫn là chặn bắt buộc — cả ba tạp chí Q2 đều "
    "yêu cầu IRB number cụ thể. NCS cần hoàn tất trong tuần đầu 6/2026."
)

add_heading("6.1. Kế hoạch dự phòng (fallback) nếu Frontiers desk-reject", level=2)

add_para(
    "Nếu Frontiers in Psychiatry desk-reject (từ chối ở vòng editor, không gửi review), "
    "em đã chuẩn bị fallback bốn cấp xếp theo IF giảm dần và khả năng accept tăng dần:"
)

add_bullet(
    "Cấp 1 — BMC Public Health (IF ≈ 4,5, Q2 Public Health) [5]: Chuyển trong 7 ngày sau "
    "desk-reject. Viết lại Title + Abstract nhấn y tế công cộng và chính sách trường học. "
    "Decision time 3–4 tháng. Đã đăng cohort thanh thiếu niên Việt Nam. APC ~2500 GBP."
)

add_bullet(
    "Cấp 2 — Frontiers in Public Health (IF ≈ 3,0, Q2 Public Health) [6]: Section Public "
    "Mental Health chấp nhận SEM cross-sectional. Decision time ≈ 3 tháng. Cùng Frontiers "
    "family nên có thể dùng Tier Transfer (không cần resubmit từ đầu). APC ~2950 USD."
)

add_bullet(
    "Cấp 3 — PLOS Mental Health (mới 2024, IF dự kiến ≈ 2,5–3,0) [12]: Scope rộng chấp "
    "nhận SEM cross-sectional. Acceptance rate ≈ 45%, decision time ≈ 2 tháng, APC "
    "~1800 USD. Nhược: chưa có IF chính thức, cần đối chiếu yêu cầu HNUE về công bố Q2."
)

add_bullet(
    "Cấp 4 — PLOS ONE (IF ≈ 3,7, broader scope) [13]: Phương án cuối. Chấp nhận mọi "
    "thiết kế đúng phương pháp. Decision time 4–6 tháng. APC ~1800 USD. Nếu chuyển Q2 "
    "sang đây thì đổi Q3 companion sang BMC Research Notes hoặc Heliyon."
)

add_para(
    "Fallback chỉ kích hoạt khi Frontiers desk-reject. Nếu Frontiers gửi review và yêu cầu "
    "major revision, tập trung response letter thay vì chuyển tạp chí — acceptance rate "
    "cuối của major revision ≈ 60–70%."
)

add_para(
    "Em chờ thầy MĐ + thầy NMĐ phê duyệt; sau khi nhận phản hồi, em triển khai bản nháp v8 "
    "SEM-only theo lịch trong mục 5."
)

# ---------------- 7. THAM KHẢO ----------------
add_heading("7. THAM KHẢO", level=1)

refs = [
    "[1] Draft_Q1_SongNgu_v7_01062026.docx — Bản nháp song ngữ Anh–Việt phiên bản 7 "
    "của bài Q1 (mixed-methods, BMC Psychiatry). Lưu trữ nội bộ tại "
    "bai-bao-Q1/Draft_Q1_SongNgu_v7_01062026.docx. Tham chiếu nội bộ.",

    "[2] 4VanDe_BLOCKING_Q1Q3_v2_01062026.docx — Tài liệu Bốn vấn đề chặn phiên bản 2 "
    "(ngày 01/06/2026). Lưu trữ nội bộ tại bai-bao-Q1/4VanDe_BLOCKING_Q1Q3_v2_01062026.docx. "
    "Tham chiếu nội bộ.",

    "[3] Braun, V., & Clarke, V. (2006). Using thematic analysis in psychology. "
    "Qualitative Research in Psychology, 3(2), 77–101. "
    "DOI: 10.1191/1478088706qp063oa.",

    "[4] Frontiers in Psychiatry. Author Guidelines & Article Types — Adolescent and "
    "Young Adult Psychiatry Section. Frontiers Media SA. "
    "Truy cập: https://www.frontiersin.org/journals/psychiatry/sections/adolescent-and-young-adult-psychiatry. "
    "Tra cứu JCR 2024 cho IF.",

    "[5] BMC Public Health. Submission Guidelines · About the Journal. Springer Nature. "
    "Truy cập: https://bmcpublichealth.biomedcentral.com/submission-guidelines. "
    "Tra cứu JCR 2024 cho IF.",

    "[6] Frontiers in Public Health. Author Guidelines — Public Mental Health Section. "
    "Frontiers Media SA. "
    "Truy cập: https://www.frontiersin.org/journals/public-health/sections/public-mental-health. "
    "Tra cứu JCR 2024 cho IF.",

    "[7] Luận án TS NCS Công Thị Hằng, phiên bản 01_LuanAn_v3_1_FixCoping_28052026.docx. "
    "Các bảng đã đối chiếu: Bảng 27, 30, 33 (yếu tố nguy cơ → phân loại lo âu); "
    "Bảng 36, 39, 42, 45 (yếu tố bảo vệ → phân loại lo âu); Bảng 47 (mô hình bậc cao "
    "YTNC + YTBV → tổng RLLA, R² = 0,598). Tham chiếu nội bộ.",

    "[8] First, M. B., Yousif, L. H., Clarke, D. E., Wang, P. S., Gogtay, N., & "
    "Appelbaum, P. S. (2022). DSM-5-TR: overview of what's new and what's changed. "
    "World Psychiatry, 21(2), 218–219. DOI: 10.1002/wps.20989. PMID: 35524596. "
    "Tham chiếu khung chẩn đoán hiện hành.",

    "[9] Pezzella, P. (2022). ICD-11 is now officially in effect worldwide. "
    "World Psychiatry, 21(2), 331–332. DOI: 10.1002/wps.20982. PMID: 35524598. "
    "Tham chiếu hệ phân loại ICD-11 hiệu lực toàn cầu 01/01/2022.",

    "[10] Web of Science. Journal Citation Reports (JCR) 2024 edition · Clarivate Analytics. "
    "Truy cập nội bộ qua thư viện HNUE / VNU. Tra cứu IF và phân hạng Q1/Q2/Q3/Q4 "
    "cho ba tạp chí ứng viên.",

    "[11] Tin nhắn chỉ đạo của thầy MĐ ngày 07/06/2026 — ba điểm chốt: bỏ phần định tính, "
    "chuyển sang thuần định lượng SEM tích hợp, chuyển Q1 → Q2. Lưu trong nhật ký trao đổi "
    "nội bộ với nhóm tác giả. Tham chiếu nội bộ.",

    "[12] PLOS Mental Health. Submission Guidelines and Author Information. Public "
    "Library of Science. Truy cập: https://journals.plos.org/mentalhealth/s/submission-guidelines. "
    "Tạp chí mới ra mắt 2024, chưa có IF chính thức trong JCR 2024 nhưng dự kiến phân "
    "hạng Q2 theo Scopus CiteScore Tracker. Dùng cho fallback Cấp 3.",

    "[13] PLOS ONE. Submission Guidelines and Author Information. Public Library of "
    "Science. Truy cập: https://journals.plos.org/plosone/s/submission-guidelines. "
    "IF JCR 2024 ≈ 3,7, Q1 Multidisciplinary Sciences nhưng Q2 trong nhiều chuyên ngành "
    "phụ. Dùng cho fallback Cấp 4.",

    "[14] Web of Science. Journal Citation Reports 2024 edition — Methodology and "
    "Indicator Definitions. Clarivate Analytics, 06/2024. "
    "Truy cập nội bộ qua thư viện HNUE / VNU. Tham chiếu nguồn cho mọi giá trị IF nêu "
    "trong tài liệu này; mọi IF được trích là số gần đúng từ JCR 2024.",

    "[15] Frontiers Media SA. Article Types — Original Research, Brief Research Report, "
    "Review. Frontiers Author Hub. "
    "Truy cập: https://www.frontiersin.org/guidelines/author-guidelines. "
    "Tham chiếu format compliance (12.000 từ, 15 hình/bảng, 75 references) và Brief "
    "Research Report (3.000 từ) cho phương án định tính follow-up.",
]
for r in refs:
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.5)
    run = p.add_run(r)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(11)

# ----------------------------------------------------------------------
# Strip metadata
# ----------------------------------------------------------------------
core = doc.core_properties
core.author = ""
core.last_modified_by = ""
core.comments = ""
core.category = ""
core.keywords = ""
core.subject = ""
core.title = ""

# Save
os.makedirs(os.path.dirname(OUT_PATH), exist_ok=True)
doc.save(OUT_PATH)

# Word count
total = 0
d = Document(OUT_PATH)
for p in d.paragraphs:
    total += len(p.text.split())
for t in d.tables:
    for row in t.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                total += len(p.text.split())
size = os.path.getsize(OUT_PATH)
print(f"SAVED: {OUT_PATH}")
print(f"SIZE: {size} bytes")
print(f"WORD COUNT: {total}")
