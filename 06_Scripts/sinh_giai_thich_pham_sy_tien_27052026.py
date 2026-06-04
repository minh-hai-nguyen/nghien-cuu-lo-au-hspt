# -*- coding: utf-8 -*-
"""Sinh doc giai thich citation Pham Sy Tien 2024 cho thay.
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '09_GiaiThich_PhamSyTien_2024_27052026.docx')

RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(220, 130, 0)
GREEN = RGBColor(0, 112, 0)


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in d.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)
    cp = d.core_properties
    cp.author = ''; cp.title = ''; cp.subject = ''
    cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
    return d


def H(d, text, level=1):
    h = d.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)


def P(d, text, bold=False, italic=False, color=None, indent=True):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def P_quote(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    r = p.add_run('"' + text + '"')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.italic = True


d = doc_init()

H(d, 'GIẢI TRÌNH VỀ TRÍCH DẪN PHẠM SỸ TIẾN VÀ CỘNG SỰ (2024)', 1)
P(d, '(NCS Công Thị Hằng — 27/05/2026)', italic=True, indent=False)
d.add_paragraph()

P(d, 'Kính gửi thầy,')
P(d, 'Em xin trình bày cách hiểu và phân tích về câu trích dẫn mà thầy hỏi.')
d.add_paragraph()

H(d, '1. Câu trích cần giải thích', 2)
P_quote(d, 'Phạm Sỹ Tiến và cộng sự (2024) khảo sát trên 546 vị thành niên tại các cơ sở bảo trợ xã hội Việt Nam ghi nhận chất lượng chăm sóc tình cảm — một thành tố của sự hỗ trợ từ cha mẹ — có hệ số β = -0,40 với điểm sức khỏe tâm thần.')
d.add_paragraph()

H(d, '2. Cách đọc nội dung khoa học của câu', 2)
P(d, 'Câu này gồm 4 thành phần thông tin chính, em xin tách ra để thầy dễ kiểm tra:')
d.add_paragraph()

P(d, 'a) Đối tượng nghiên cứu:', bold=True, indent=False)
P(d, '546 vị thành niên đang sống tại các cơ sở bảo trợ xã hội ở Việt Nam — tức là nhóm trẻ vị thành niên đặc thù, không sống cùng cha mẹ ruột mà được nuôi dưỡng tại trung tâm bảo trợ/mái ấm/cơ sở chăm sóc thay thế của nhà nước hoặc tổ chức xã hội. Đây là nhóm yếu thế, có nguy cơ cao về sức khỏe tâm thần.')
d.add_paragraph()

P(d, 'b) Biến độc lập được khảo sát:', bold=True, indent=False)
P(d, '"Chất lượng chăm sóc tình cảm" (emotional care quality) — được định nghĩa là một thành tố/khía cạnh của khái niệm rộng hơn là "sự hỗ trợ từ cha mẹ" (parental support). Trong nhóm trẻ tại cơ sở bảo trợ, người làm "vai trò cha mẹ" có thể là cán bộ chăm sóc/người giám hộ thay thế, chứ không phải cha mẹ ruột — đây là điểm cần lưu ý về tính khái quát hóa kết quả.')
d.add_paragraph()

P(d, 'c) Biến phụ thuộc:', bold=True, indent=False)
P(d, '"Điểm sức khỏe tâm thần" — đây là biến tổng hợp, có thể tính từ thang đo như SDQ, DASS-21, hoặc thang chuyên dụng cho trẻ vị thành niên trong môi trường thay thế gia đình. Điểm cao thường biểu thị "vấn đề nhiều" (hoặc ngược lại, tùy hướng tính của thang).')
d.add_paragraph()

P(d, 'd) Hệ số chuẩn hóa β = -0,40:', bold=True, indent=False)
P(d, 'Đây là hệ số hồi quy chuẩn hóa (standardized beta coefficient) trong mô hình hồi quy tuyến tính đa biến hoặc mô hình phương trình cấu trúc (SEM). Diễn giải:')
P(d, '• Dấu âm (-) → hai biến có mối liên hệ NGHỊCH chiều: khi chất lượng chăm sóc tình cảm tăng thì điểm sức khỏe tâm thần (theo nghĩa "vấn đề SKTT") giảm.')
P(d, '• Độ lớn 0,40 → theo quy ước Cohen (1988), đây là cỡ hiệu ứng MẠNH (large effect). Mức 0,1 nhỏ, 0,3 trung bình, 0,5 lớn → 0,40 nằm ở khoảng trung bình–lớn.')
P(d, '• Cách diễn giải đơn giản: trong nhóm trẻ tại cơ sở bảo trợ, chất lượng chăm sóc tình cảm là yếu tố BẢO VỆ mạnh đối với sức khỏe tâm thần. Cứ tăng 1 độ lệch chuẩn (SD) của chăm sóc tình cảm thì điểm vấn đề SKTT giảm 0,40 SD.')
d.add_paragraph()

H(d, '3. Ý nghĩa của trích dẫn này trong luận án', 2)
P(d, 'Trích dẫn được đưa vào phần bàn luận về "yếu tố bảo vệ — sự hỗ trợ từ cha mẹ" để khẳng định:')
P(d, '• Chất lượng chăm sóc tình cảm là một thành tố quan trọng của khái niệm "sự hỗ trợ từ cha mẹ"; nó có thể được đo lường tách biệt và có hệ số tác động riêng.')
P(d, '• Phát hiện ở nhóm trẻ tại cơ sở bảo trợ — vốn đã thiếu hụt cha mẹ ruột — vẫn cho thấy sức mạnh của chăm sóc tình cảm. Điều này hỗ trợ giả thuyết: chăm sóc tình cảm CÓ TÍNH HÀM Ý CHUNG vượt ra ngoài nhóm gia đình ruột thịt, và áp dụng được cả trong môi trường thay thế.')
P(d, '• Từ đó, trong luận án của em, khi xây dựng khung "sự hỗ trợ từ cha mẹ" cho học sinh THCS, việc tách biến "chất lượng chăm sóc tình cảm" làm một biến độc lập đo riêng (chứ không gộp với toàn bộ "hỗ trợ cha mẹ") là phù hợp với hướng nghiên cứu của Phạm Sỹ Tiến và cộng sự (2024).')
d.add_paragraph()

H(d, '4. Lưu ý về tính so sánh được giữa hai mẫu', 2)
P(d, 'Em xin thưa với thầy có một số khác biệt giữa nghiên cứu của Phạm Sỹ Tiến (2024) và luận án của em mà cần được nói rõ khi trích dẫn:')
P(d, '• Phạm Sỹ Tiến nghiên cứu trên trẻ tại cơ sở bảo trợ; em nghiên cứu trên học sinh THCS sống tại gia đình.')
P(d, '• Người "thực hiện chăm sóc tình cảm" trong nghiên cứu của Phạm Sỹ Tiến có thể là cán bộ trung tâm thay vì cha mẹ ruột.')
P(d, '• Hệ số β = -0,40 ở nhóm trẻ tại cơ sở bảo trợ có thể lớn hơn hoặc nhỏ hơn ở nhóm học sinh sống cùng cha mẹ ruột (do điểm xuất phát ban đầu khác nhau).')
P(d, 'Vì vậy, khi đưa β này vào luận án, em chỉ nên dùng như một bằng chứng định hướng (directional evidence), KHÔNG dùng làm con số đối chiếu trực tiếp với hệ số em sẽ tính được trên mẫu THCS.')
d.add_paragraph()

# Caveat
H(d, '5. Lưu ý về nguồn tài liệu', 2)
P(d, 'Trước khi nộp bản luận án cuối cho thầy, em xin xác nhận lại đây là nguồn tham khảo mà em sẽ verify đầy đủ:', italic=True, color=ORANGE)
P(d, '• Bài báo gốc: cần kiểm tra DOI, tạp chí, năm xuất bản và truy cập bản PDF/full text.')
P(d, '• Dữ liệu cụ thể (n=546, β=-0,40, biến độc lập, biến phụ thuộc, công cụ đo): cần đối chiếu với bảng kết quả trong PDF gốc, KHÔNG dùng số liệu chỉ qua tóm tắt.')
P(d, '• Nếu bài báo có sai khác về N, hệ số β, hoặc khung khái niệm, em sẽ cập nhật ngay trong bản luận án và báo cáo lại thầy.')
d.add_paragraph()

P(d, 'Em sẽ verify nguồn này vào ngày làm việc tiếp theo và gửi lại thầy bản trích dẫn đã đối chiếu đầy đủ.', italic=True, color=ORANGE)
d.add_paragraph()

P(d, 'Em xin cảm ơn thầy.', indent=False)
P(d, 'Trân trọng,', indent=False)
P(d, '', indent=False)
P(d, 'NCS Công Thị Hằng', bold=True, indent=False)

# Remove watermark + header/footer
for sec in d.sections:
    for hf in [sec.header, sec.first_page_header, sec.even_page_header,
               sec.footer, sec.first_page_footer, sec.even_page_footer]:
        for elem in list(hf._element.iter()):
            if elem.tag in (qn('w:pict'), qn('w:object')):
                if elem.getparent() is not None:
                    elem.getparent().remove(elem)
        for p in hf.paragraphs:
            for r in p.runs:
                r.text = ''

d.save(OUT)
print(f"Da luu: {OUT}")
from docx import Document as D
dd = D(OUT)
w = sum(len(p.text.split()) for p in dd.paragraphs)
print(f"Paragraphs: {len(dd.paragraphs)}, Words: ~{w}, Size: {os.path.getsize(OUT)//1024}KB")
