# -*- coding: utf-8 -*-
"""Verify LA v3_3 - check remaining Xu 2022 + Sakia."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_3_FixSaikiaXu_27052026.docx')
d = Document(LA)

# Check Xu year 2022 remaining
print("=== Remaining 'Xu' + '2022' near each other ===")
for i, p in enumerate(d.paragraphs):
    if 'Xu' in p.text:
        # Check if '2022' near 'Xu'
        for kw in ['Xu', 'Qingqing']:
            pos = 0
            while True:
                pos = p.text.find(kw, pos)
                if pos == -1: break
                # window of 30 chars after
                ctx = p.text[pos:pos+80]
                if '2022' in ctx:
                    print(f"  para {i}: ...{ctx}...")
                pos += 1
print("\n=== Remaining 'Sakia' (spelling typo) ===")
for i, p in enumerate(d.paragraphs):
    if 'Sakia' in p.text:
        print(f"  para {i}: {p.text[:200]}")

print("\n=== '38,42%' context ===")
for i, p in enumerate(d.paragraphs):
    if '38,42%' in p.text:
        print(f"  para {i}: {p.text[:400]}")

print("\n=== Para 238 (Xu intro) - full ===")
print(d.paragraphs[238].text[:500])
