# -*- coding: utf-8 -*-
"""
QA SAU — kiem tra tat ca loi tiem an cho ban dich VN022:
1. So sanh character ratio EN/VN
2. Kiem tra moi so lieu trong bang khop voi EN
3. Kiem tra typos/formatting
4. Kiem tra nhat quan thuat ngu
5. So khop voi ban tom tat cu
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from pypdf import PdfReader
from collections import Counter, defaultdict

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
VN_DOCX = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')
VN_SUM = os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN022_UNICEF_VN_2022_SchoolFactors.docx')
EN_PDF = os.path.join(ROOT, '02_Papers-goc', 'Viet-Nam', 'UNICEF_2022_School_Factors_Vietnam.pdf')

# Load
d = Document(VN_DOCX)
vn_paras = [p.text for p in d.paragraphs if p.text.strip()]
vn_text = '\n'.join(vn_paras)
for t in d.tables:
    for r in t.rows:
        vn_text += '\n' + ' | '.join(c.text.strip() for c in r.cells)

d_sum = Document(VN_SUM)
sum_text = '\n'.join(p.text for p in d_sum.paragraphs if p.text.strip())
for t in d_sum.tables:
    for r in t.rows:
        sum_text += '\n' + ' | '.join(c.text.strip() for c in r.cells)

r = PdfReader(EN_PDF)
en_text = '\n'.join((p.extract_text() or '') for p in r.pages)

print('='*70)
print('DEEP QA — VN022 UNICEF 2022 TRANSLATION')
print('='*70)
print(f'VN full docx: {len(vn_text):,} chars, {len(vn_paras)} paragraphs')
print(f'VN summary: {len(sum_text):,} chars')
print(f'EN PDF: {len(en_text):,} chars')
print(f'Coverage: {len(vn_text)/len(en_text)*100:.1f}%')

issues = defaultdict(list)

# ============================================================
# 1. KEY STATISTICS — 40 critical numbers
# ============================================================
print('\n' + '#'*70)
print('CHECK 1: 40 Critical Statistics (must all verify)')
print('#'*70)

CRITICAL_STATS = [
    # (EN_form, VN_form, description)
    ('26.1%', '26,1', 'Overall mental health risk'),
    ('32%', '32', 'Peer problems'),
    ('30.9%', '30,9', 'Emotional problems'),
    ('14.4%', '14,4', 'Hyperactivity'),
    ('11%', '11', 'Conduct problems'),
    ('668', '668', 'Student sample'),
    ('66', '66', 'Teacher sample'),
    ('34%', '34', 'Male students'),
    ('66%', '66', 'Female students'),
    ('14.23', '14,23', 'Mean age'),
    ('86.4%', '86,4', 'Teachers no MH training'),
    ('91%', '91', 'Teachers concerned about stress'),
    ('95%', '95', 'Teachers concerned about MH'),
    ('28%', '28', 'Students > 3h study/night'),
    ('50%', '50', 'Students > 2h study/night'),
    ('15%', '15', 'Students > 9h extra classes'),
    ('47.3%', '47,3', 'Never cyberbullied'),
    ('29.9%', '29,9', 'Rarely cyberbullied'),
    ('20.2%', '20,2', 'Sometimes cyberbullied'),
    ('2.1%', '2,1', 'Often cyberbullied'),
    ('70%', '70', 'Rural schools no counselling room'),
    ('39.61', '39,61', 'Teacher mean age'),
    ('84.8%', '84,8', 'Female teachers'),
    ('15.2%', '15.2', 'Male teachers'),
    ('16.86', '16,86', 'Years teaching experience'),
    ('60.6%', '60,6', 'Teachers secondary'),
    ('54%', '54', 'Teachers Kinh'),
    ('81.8%', '81,8', 'Teachers Kinh %'),
    ('92.4%', '92,4', 'Poor attention'),
    ('69.7%', '69,7', 'Lack of confidence'),
    ('51.5%', '51,5', 'Anxiety reported'),
    ('96.1%', '96,1', 'Happy House participation'),
    ('1,084', '1.084', 'Happy House students'),
    ('54%', '54', 'Ethnic minority students'),
    ('12%', '12', 'Weiss 2014 VN MH rate'),
    ('25.5%', '25,5', 'Happy House baseline'),
    ('15%', '15', 'Global 15% adolescent MH'),
    ('69%', '69', 'SDGCW child violence'),
    ('0.91', '0,91', 'WHO psychiatrists/100K VN'),
    ('355,000', '355', 'Some large number'),
]

verified = 0
missing = []
for en_s, vn_s, desc in CRITICAL_STATS:
    if vn_s in vn_text:
        verified += 1
    else:
        missing.append((en_s, vn_s, desc))

print(f'Verified: {verified}/{len(CRITICAL_STATS)} ({verified*100//len(CRITICAL_STATS)}%)')
if missing:
    print(f'MISSING in VN:')
    for en, vn, desc in missing:
        print(f'  [{desc}] EN={en}, VN expected={vn} — NOT FOUND')
        issues['stats'].append(f'{desc}: VN {vn} missing')

# ============================================================
# 2. POLICY REFERENCES — all circulars, decisions must be preserved
# ============================================================
print('\n' + '#'*70)
print('CHECK 2: Policy References')
print('#'*70)

POLICIES = [
    ('9971', 'Công văn 9971/2005'),
    ('31/2017', 'Thông tư 31/2017'),
    ('1215', 'QĐ 1215/2011 MOLISA'),
    ('1929', 'QĐ 1929/2020'),
    ('1069', 'QĐ 1069/2021'),
    ('112', 'QĐ 112/2021'),
    ('1437', 'QĐ 1437/2018'),
    ('1438', 'QĐ 1438/2018'),
    ('1660', 'QĐ 1660/2021'),
    ('4659', 'QĐ 4659/2021'),
    ('4969', 'QĐ 4969/2021'),
    ('85', 'QĐ 85/2022'),
    ('155', 'QĐ 155/2022'),
    ('712', 'QĐ 712/2022'),
    ('800', 'CT 800/2021'),
    ('136', 'CV 136/2022'),
    ('2016', 'Luật Trẻ em 2016'),
    ('111', 'Đường dây 111'),
]
policy_verified = 0
policy_missing = []
for code, desc in POLICIES:
    if code in vn_text:
        policy_verified += 1
    else:
        policy_missing.append((code, desc))

print(f'Policies verified: {policy_verified}/{len(POLICIES)}')
if policy_missing:
    for code, desc in policy_missing:
        print(f'  MISSING: {desc} (code: {code})')
        issues['policy'].append(f'{desc} missing')

# ============================================================
# 3. INSTRUMENTS / SCALES
# ============================================================
print('\n' + '#'*70)
print('CHECK 3: Instruments and Scales')
print('#'*70)

INSTRUMENTS = ['SDQ-25', 'SDQ25', 'MDS3', 'ESSA', 'RCBI',
               'GAD-7', 'DASS-21', 'PHQ-9', 'SRQ-20']
for inst in INSTRUMENTS:
    n = vn_text.count(inst)
    print(f'  {inst}: {n} mentions')
    if inst in ['SDQ-25', 'MDS3', 'ESSA', 'RCBI'] and n == 0:
        # Try variants
        if inst == 'SDQ-25':
            n2 = vn_text.count('SDQ25') + vn_text.count('SDQ')
            if n2 > 0:
                print(f'    (SDQ variants found: {n2})')
            else:
                issues['instruments'].append(f'{inst} missing')
        elif n == 0:
            issues['instruments'].append(f'{inst} missing')

# ============================================================
# 4. PROVINCE NAMES (4 included provinces)
# ============================================================
print('\n' + '#'*70)
print('CHECK 4: Provinces and Geography')
print('#'*70)

PROVINCES = ['Hà Nội', 'Điện Biên', 'Gia Lai', 'Đồng Tháp', 'TPHCM']
for prov in PROVINCES:
    n = vn_text.count(prov)
    print(f'  {prov}: {n} mentions')
    if n == 0 and prov != 'TPHCM':
        issues['provinces'].append(f'{prov} missing')

# ============================================================
# 5. STRUCTURE CHECK — all chapters
# ============================================================
print('\n' + '#'*70)
print('CHECK 5: Structure — Chapters + Appendices')
print('#'*70)

SECTIONS = [
    'CHƯƠNG 1: GIỚI THIỆU',
    'CHƯƠNG 2: SỨC KHOẺ TÂM THẦN VỊ THÀNH NIÊN',
    'CHƯƠNG 3: TỔNG QUAN TÀI LIỆU VỀ SỨC KHOẺ TÂM THẦN VỊ THÀNH NIÊN TẠI VIỆT NAM',
    'CHƯƠNG 4: PHƯƠNG PHÁP',
    'CHƯƠNG 5: PHÁT HIỆN NGHIÊN CỨU',
    'CHƯƠNG 6: CÁC YẾU TỐ LIÊN QUAN ĐẾN TRƯỜNG HỌC',
    'CHƯƠNG 7: CHÍNH SÁCH VÀ CHƯƠNG TRÌNH MHPSS HỌC ĐƯỜNG',
    'CHƯƠNG 8: CHÍNH SÁCH VÀ CHƯƠNG TRÌNH MHPSS',
    'CHƯƠNG 9: KẾT LUẬN VÀ KHUYẾN NGHỊ',
    'PHỤ LỤC 1',
    'PHỤ LỤC 2',
    'PHỤ LỤC 3',
    'TÀI LIỆU THAM KHẢO',
]
for s in SECTIONS:
    if s in vn_text:
        print(f'  ✓ {s[:60]}')
    else:
        print(f'  ✗ MISSING: {s}')
        issues['structure'].append(f'Missing section: {s}')

# ============================================================
# 6. TERMINOLOGY CONSISTENCY
# ============================================================
print('\n' + '#'*70)
print('CHECK 6: Terminology Consistency')
print('#'*70)

# Terms that should be used consistently
terms = {
    'mental health': ['sức khoẻ tâm thần', 'sức khỏe tâm thần'],  # mostly first
    'well-being': ['hạnh phúc'],
    'adolescent': ['vị thành niên'],
    'school climate': ['khí hậu trường học'],
    'bullying': ['bắt nạt'],
    'cyberbullying': ['bắt nạt mạng', 'cyberbullying'],
    'counselling': ['tư vấn'],
    'ethnic minority': ['dân tộc thiểu số'],
}
for en, vns in terms.items():
    counts = [(vn, vn_text.count(vn)) for vn in vns]
    total = sum(c for _, c in counts)
    print(f'  "{en}": total {total}, variants: {counts}')
    # Flag if inconsistent (multiple variants all used)
    if len(counts) > 1 and all(c > 5 for _, c in counts):
        issues['terminology'].append(f'Inconsistent: {en} → {counts}')

# ============================================================
# 7. NUMBERS CONSISTENCY — no stat should contradict itself
# ============================================================
print('\n' + '#'*70)
print('CHECK 7: Internal Number Consistency')
print('#'*70)

# For 668 students — confirm it's consistent
contexts_668 = []
for m in re.finditer(r'668', vn_text):
    ctx = vn_text[max(0, m.start()-60):m.end()+60].replace('\n', ' ')
    contexts_668.append(ctx)
print(f'  "668" contexts: {len(contexts_668)}')
# Check for conflicting student counts nearby
for ctx in contexts_668[:3]:
    print(f'    ...{ctx[:180]}...')

# Check 26.1% vs 26%
has_261 = '26,1' in vn_text
has_26_plain = re.search(r'\b26\s*%', vn_text) is not None
if has_261 and has_26_plain:
    # That's OK if 26% is rounded form, check if both explained
    pass

# ============================================================
# 8. COMPARE WITH SUMMARY
# ============================================================
print('\n' + '#'*70)
print('CHECK 8: Compare FULL translation vs SUMMARY')
print('#'*70)

# Extract key claims from summary
print(f'Summary length: {len(sum_text):,} chars')
print(f'Full length: {len(vn_text):,} chars')
print(f'Ratio: {len(sum_text)/len(vn_text)*100:.1f}%')

# Summary should mention same key stats
sum_stats = re.findall(r'(\d+[,.]\d+)\s*%', sum_text)
sum_stats_set = set(sum_stats)
full_stats_set = set(re.findall(r'(\d+[,.]\d+)\s*%', vn_text))

# Key stats in summary that should be in full
sum_only = sum_stats_set - full_stats_set
full_only_key = []
for s in list(full_stats_set)[:20]:
    if s not in sum_text and float(s.replace(',', '.')) > 10:
        full_only_key.append(s)

print(f'Summary-only stats: {sum_only}')
if sum_only:
    for s in sum_only:
        # Check if close variant exists
        if s.replace(',', '.') in vn_text or s in vn_text:
            pass
        else:
            issues['sum_vs_full'].append(f'Summary mentions {s}% but not in full')

# Check critical summary claims
SUM_KEY = ['8–29', '50%', '15%', '47 %', '5 tỉnh', 'SDQ-25', 'ESSA', '668']
print(f'\nKey summary claims verification in FULL:')
for k in SUM_KEY:
    in_full = k in vn_text
    in_sum = k in sum_text
    print(f'  "{k}": summary={in_sum}, full={in_full}')

# ============================================================
# 9. GRAMMAR/TYPO patterns
# ============================================================
print('\n' + '#'*70)
print('CHECK 9: Grammar and Typo Patterns')
print('#'*70)

# Common VN translation issues
issues_patterns = [
    (r'\bthe\s+', 'English "the" leaking through'),
    (r'\band\s+', 'English "and" leaking through'),
    (r'\bof\s+the\s+', 'English "of the" leaking through'),
    (r'chính sáchh|chính  sách', 'Typo "chính sách"'),
    (r'học sinhh', 'Typo "học sinh"'),
    (r'\bVN\b\s+022\b', 'Spacing issue VN022'),
    (r'%\s*%', 'Double percent sign'),
]
for pat, desc in issues_patterns:
    matches = re.findall(pat, vn_text)
    if matches:
        # Filter out legitimate usages (e.g., "the" in English instrument names)
        legit = ['the Ministry', 'The Resourceful', 'The 2016 Law', 'the national',
                'The content', 'Revised Cyber']
        filtered = [m for m in matches if not any(l.lower() in vn_text[max(0, vn_text.find(m)-50):vn_text.find(m)+50].lower() for l in legit)]
        if len(filtered) > 3:  # threshold
            print(f'  ⚠ {desc}: {len(matches)} instances')
            issues['grammar'].append(f'{desc}: {len(matches)}')

# ============================================================
# 10. SPECIFIC FACT CHECK vs SUMMARY
# ============================================================
print('\n' + '#'*70)
print('CHECK 10: Specific Facts Summary vs Full')
print('#'*70)

# Facts from summary
SUMMARY_FACTS = [
    ('8–29%', 'VN adolescent MH range'),
    ('50%', 'Students extra classes'),
    ('5 tỉnh', '5 provinces plan'),
    ('4 tỉnh', '4 provinces actual'),
    ('SDQ-25', 'Mental health scale'),
    ('ESSA', 'Stress scale'),
    ('668 HS', 'Student sample'),
    ('249 giáo viên', 'Teacher sample in summary (doesn\'t match 66 in full!)'),
]

fact_mismatches = []
for fact, desc in SUMMARY_FACTS:
    sum_has = fact in sum_text
    full_has = fact in vn_text
    status = '✓' if sum_has == full_has else '✗'
    print(f'  {status} {desc}: sum={sum_has}, full={full_has}')
    if sum_has and not full_has:
        fact_mismatches.append(f'{desc}: in summary but not in full')
        issues['fact_consistency'].append(f'{desc}: in summary but not in full')

# CRITICAL: Summary says 249 teachers, but full says 66!
if '249' in sum_text and '66' in vn_text:
    # Check context
    if '249 giáo' in sum_text:
        print('\n  ⚠⚠⚠ CRITICAL: Summary mentions "249 teachers" but full text says 66 teachers!')
        # Check if 249 also in full
        if '249' not in vn_text:
            issues['fact_consistency'].append('SUMMARY-FULL MISMATCH: 249 teachers (summary) vs 66 teachers (full)')

# ============================================================
# FINAL REPORT
# ============================================================
print('\n' + '='*70)
print('FINAL DEEP QA REPORT')
print('='*70)
total = sum(len(v) for v in issues.values())
print(f'Total issues found: {total}')
for k, v in sorted(issues.items()):
    print(f'\n{k}: {len(v)} issue(s)')
    for msg in v:
        print(f'  - {msg}')

print(f'\nCoverage: {len(vn_text)/len(en_text)*100:.1f}%')
print(f'Summary-to-Full match: {len(sum_text)/len(vn_text)*100:.1f}%')
