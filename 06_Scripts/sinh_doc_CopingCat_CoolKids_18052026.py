# -*- coding: utf-8 -*-
"""Sinh doc tra loi: Coping Cat & Cool Kids - tac gia, nam ra doi. 18/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Tra_loi_CopingCat_CoolKids_18052026.docx')
PAGE_W = 15.5

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)

def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    return p

def bullet(doc, text, bold_prefix=None, size=12):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = True
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    return p

def add_hyperlink(paragraph, url, text, size=11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def ref_item(doc, text, doi_label=None, doi_url=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    if doi_label and doi_url:
        sep = p.add_run('  '); sep.font.name = 'Times New Roman'; sep.font.size = Pt(11)
        add_hyperlink(p, doi_url, doi_label, size=11)

def table(doc, headers, rows, widths, hdr_size=10, body_size=10):
    assert abs(sum(widths) - PAGE_W) < 0.1, sum(widths)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(hdr_size)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(body_size)
                    if ci == 0: r.bold = True
    return t

# ============================================================
doc = make_doc()

ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('TRẢ LỜI: HAI CHƯƠNG TRÌNH "COPING CAT" VÀ "COOL KIDS"\n'
                'RA ĐỜI NĂM NÀO, DO AI XÂY DỰNG?')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tài liệu trả lời câu hỏi — Ngày 18/05/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

P(doc, 'Câu hỏi: Hai chương trình "Coping Cat" và "Cool Kids" ra đời từ năm nào, ai là tác '
       'giả?', italic=True, size=11)

# ---- 1. Coping Cat ----
H(doc, '1. Chương trình "Coping Cat"', 1)
H(doc, '1.1. Tác giả và năm ra đời', 2)
P(doc, '"Coping Cat" (tạm dịch: "Chú mèo biết ứng phó") do nhà tâm lý học '
       'Philip C. Kendall cùng cộng sự xây dựng tại Phòng khám Rối loạn Lo âu Trẻ em và Vị '
       'thành niên, Đại học Temple (Hoa Kỳ). Đây là chương trình mở đường cho việc áp dụng '
       'liệu pháp nhận thức – hành vi (Cognitive Behavioral Therapy — CBT) vào điều trị lo âu '
       'ở trẻ em.')
P(doc, 'Về mốc thời gian: sổ tay thực hành dành cho trẻ ("Coping Cat Workbook") được Kendall '
       'xuất bản lần đầu năm 1990, đánh dấu sự ra đời của chương trình. Bằng chứng khoa học '
       'nền tảng đến từ năm 1994, khi Kendall công bố thử nghiệm lâm sàng ngẫu nhiên (RCT) đầu '
       'tiên trên thế giới chứng minh hiệu quả của CBT đối với trẻ lo âu (trên 47 trẻ, so '
       'sánh với nhóm chờ điều trị). Có thể tóm tắt: chương trình hình thành đầu thập niên '
       '1990, được kiểm chứng chính thức năm 1994.')

H(doc, '1.2. Nội dung và đặc điểm', 2)
P(doc, '"Coping Cat" là một chương trình CBT 16 buổi, theo cẩm nang chuẩn hóa, dành cho trẻ '
       '7–13 tuổi mắc rối loạn lo âu lan tỏa, rối loạn lo âu xã hội hoặc rối loạn lo âu chia '
       'ly. Chương trình gồm hai nửa: nửa đầu rèn kỹ năng (nhận diện dấu hiệu lo âu, thư giãn, '
       'nhận ra và điều chỉnh suy nghĩ lo lắng, giải quyết vấn đề, tự khen thưởng); nửa sau là '
       'thực hành tiếp xúc dần với tình huống gây sợ. Chương trình được triển khai theo hình '
       'thức trị liệu cá nhân, có cẩm nang cho nhà trị liệu và sổ tay riêng cho trẻ. Đến nay, '
       '"Coping Cat" vẫn là một trong những chương trình CBT cho trẻ lo âu được nghiên cứu và '
       'sử dụng rộng rãi nhất trên thế giới.')

# ---- 2. Cool Kids ----
H(doc, '2. Chương trình "Cool Kids"', 1)
H(doc, '2.1. Tác giả và năm ra đời', 2)
P(doc, '"Cool Kids" do Giáo sư Ronald M. Rapee cùng cộng sự xây dựng tại Trung tâm Sức khỏe '
       'Cảm xúc (Centre for Emotional Health — nay là Lifespan Health and Wellbeing Research '
       'Centre), Đại học Macquarie, Sydney, Úc.')
P(doc, 'Về mốc thời gian: theo chính Đại học Macquarie, đến khoảng năm 2000 — sau khi Giáo sư '
       'Rapee chuyển về Macquarie — chương trình mới định hình với tên gọi "Cool Kids". Vì '
       'vậy có thể xem năm 2000 là mốc ra đời của chương trình; các cẩm nang về sau được cập '
       'nhật và tái bản nhiều lần.')

H(doc, '2.2. Mối liên hệ với "Coping Cat" — điểm quan trọng', 2)
P(doc, '"Cool Kids" không phải là một chương trình hoàn toàn độc lập, mà kế thừa trực tiếp từ '
       '"Coping Cat". Tuyến phát triển như sau: từ "Coping Cat" của Kendall (Hoa Kỳ), tác giả '
       'Paula Barrett đã soạn một bản thích ứng cho nước Úc mang tên "Coping Koala" (Sổ tay '
       'Chú gấu Koala biết ứng phó, bản thảo chưa xuất bản, Đại học Griffith, 1995). Bản '
       '"Coping Koala" này được dùng trong thử nghiệm có đối chứng nổi tiếng của Barrett, '
       'Dadds và Rapee (1996). Từ nền tảng đó, Giáo sư Rapee tiếp tục phát triển và định hình '
       'thành "Cool Kids" tại Đại học Macquarie vào khoảng năm 2000.')

H(doc, '2.3. Nội dung và đặc điểm', 2)
P(doc, '"Cool Kids" dành cho trẻ và vị thành niên 7–17 tuổi. So với "Coping Cat", "Cool Kids" '
       'thường được triển khai theo hình thức trị liệu nhóm và đặc biệt nhấn mạnh sự tham gia '
       'của cha mẹ trong quá trình can thiệp. Các thành phần CBT cốt lõi gồm: giáo dục tâm lý, '
       'tập suy nghĩ thực tế (tái cấu trúc nhận thức), tiếp xúc dần với tình huống gây lo, rèn '
       'kỹ năng xã hội và hướng dẫn cha mẹ cách hỗ trợ con. Hiện chương trình đã được dịch ra '
       'hơn 25 ngôn ngữ và sử dụng tại hơn 30 quốc gia.')

H(doc, '2.4. Phân biệt với "Cool Little Kids"', 2)
P(doc, 'Cần phân biệt "Cool Kids" với một chương trình khác cùng nhóm Macquarie là "Cool '
       'Little Kids". "Cool Little Kids" là chương trình giáo dục dành cho cha mẹ, nhằm phòng '
       'ngừa sớm lo âu ở trẻ mầm non có khí chất rụt rè – ức chế; chương trình này được công '
       'bố qua nghiên cứu của Rapee và cộng sự năm 2005. Đây là một chương trình riêng, không '
       'đồng nhất với "Cool Kids".')

# ---- 3. Bang tom tat ----
H(doc, '3. Bảng tóm tắt so sánh', 1)
table(doc,
    ['Tiêu chí', 'Coping Cat', 'Cool Kids'],
    [
        ['Tác giả chính', 'Philip C. Kendall và cộng sự',
         'Ronald M. Rapee và cộng sự'],
        ['Đơn vị', 'Đại học Temple (Hoa Kỳ)',
         'Đại học Macquarie — Trung tâm Sức khỏe Cảm xúc (Sydney, Úc)'],
        ['Năm ra đời', 'Sổ tay đầu tiên: 1990; thử nghiệm kiểm chứng kinh điển: 1994',
         'Định hình với tên "Cool Kids": khoảng năm 2000'],
        ['Đối tượng', 'Trẻ 7–13 tuổi mắc rối loạn lo âu',
         'Trẻ và vị thành niên 7–17 tuổi'],
        ['Hình thức gốc', 'Trị liệu cá nhân, 16 buổi, theo cẩm nang chuẩn hóa',
         'Trị liệu nhóm, nhấn mạnh sự tham gia của cha mẹ'],
        ['Quan hệ giữa hai chương trình',
         'Chương trình gốc, mở đường cho CBT trẻ lo âu',
         'Kế thừa từ "Coping Cat" qua bản thích ứng Úc "Coping Koala"'],
    ],
    widths=[3.3, 5.6, 6.6])

# ---- 4. Lien he bai2 ----
H(doc, '4. Liên hệ với bản thảo bai2', 1)
P(doc, 'Hai chương trình này được nhắc tới trong bản thảo bai2 (Mục 1 và Mục 2). Câu ở Mục 2 '
       '— "các giao thức Coping Cat phát triển từ thập niên 1990" — là chính xác. Tuy nhiên, '
       'danh mục tài liệu tham khảo của bai2 hiện CHƯA có tài liệu gốc nào cho "Coping Cat" '
       'và "Cool Kids". Nếu giữ hai chương trình này trong bài, nên bổ sung ít nhất: Kendall '
       '(1994) cho "Coping Cat"; Barrett, Dadds và Rapee (1996) cho tuyến "Coping Koala" dẫn '
       'tới "Cool Kids". Các tài liệu này được liệt kê đầy đủ ở Mục 5.')

# ---- 5. Tham khao ----
H(doc, '5. Tham khảo & truy vết', 1)
P(doc, 'Nguồn học thuật (đã đối chiếu để xác minh):', bold=True, size=11)
ref_item(doc,
    'Kendall, P. C. (1990). Coping Cat Workbook. Ardmore, PA: Workbook Publishing. '
    '— Ấn bản đầu tiên của sổ tay "Coping Cat".')
ref_item(doc,
    'Kendall, P. C. (1994). "Treating anxiety disorders in children: Results of a randomized '
    'clinical trial." Journal of Consulting and Clinical Psychology, 62(1), 100–110. '
    'PMID: 8034812. — Thử nghiệm lâm sàng ngẫu nhiên đầu tiên kiểm chứng CBT cho trẻ lo âu.',
    'doi:10.1037/0022-006X.62.1.100',
    'https://doi.org/10.1037/0022-006X.62.1.100')
ref_item(doc,
    'Barrett, P. M. (1995). Coping Koala Workbook. Bản thảo chưa xuất bản, School of Applied '
    'Psychology, Griffith University, Australia. — Bản thích ứng "Coping Cat" cho nước Úc, '
    'tiền thân của "Cool Kids".')
ref_item(doc,
    'Barrett, P. M., Dadds, M. R., & Rapee, R. M. (1996). "Family treatment of childhood '
    'anxiety: A controlled trial." Journal of Consulting and Clinical Psychology, 64(2), '
    '333–342. PMID: 8871418. — Thử nghiệm dùng bản "Coping Koala", nền tảng dẫn tới '
    '"Cool Kids".',
    'doi:10.1037/0022-006X.64.2.333',
    'https://doi.org/10.1037/0022-006X.64.2.333')
ref_item(doc,
    'Rapee, R. M., Kennedy, S., Ingram, M., và cộng sự (2005). "Prevention and early '
    'intervention of anxiety disorders in inhibited preschool children." Journal of '
    'Consulting and Clinical Psychology, 73(3), 488–497. PMID: 15982146. — Công bố chương '
    'trình "Cool Little Kids" (khác với "Cool Kids").',
    'doi:10.1037/0022-006X.73.3.488',
    'https://doi.org/10.1037/0022-006X.73.3.488')

P(doc, 'Nguồn thông tin chương trình (trang chính thức):', bold=True, size=11)
p = doc.add_paragraph(style='List Bullet'); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Trang chương trình "Cool Kids", Đại học Macquarie: ')
r.font.name = 'Times New Roman'; r.font.size = Pt(11)
add_hyperlink(p, 'https://coolkids.org.au/site/about', 'coolkids.org.au', size=11)
p = doc.add_paragraph(style='List Bullet'); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
r = p.add_run('Hồ sơ Giáo sư Ronald M. Rapee, Đại học Macquarie: ')
r.font.name = 'Times New Roman'; r.font.size = Pt(11)
add_hyperlink(p, 'https://researchers.mq.edu.au/en/persons/ron-rapee/',
              'researchers.mq.edu.au', size=11)

P(doc, 'Truy vết nội bộ:', bold=True, size=11)
bullet(doc, '06_Scripts/build_bai2_KHGDVN.py — bản thảo bai2 nhắc tới "Coping Cat" và "Cool '
            'Kids" ở Mục 1 (biến SECTION_1_INTRO) và Mục 2 (biến SECTION_2_METHOD).', size=11)

P(doc, 'Ghi chú: năm bảo vệ một số mốc (ví dụ thời điểm xuất bản cẩm nang "Cool Kids" qua '
       'từng phiên bản) có thể chênh lệch nhỏ giữa các nguồn; báo cáo này lấy mốc theo nguồn '
       'chính thức của Đại học Macquarie và các bài báo gốc trên PubMed.', italic=True, size=11)

doc.save(OUT)
print('DA LUU:', OUT)
print('So bang:', len(doc.tables), '| So doan:', len(doc.paragraphs))
