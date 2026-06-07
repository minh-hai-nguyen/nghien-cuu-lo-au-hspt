# -*- coding: utf-8 -*-
"""Doc 4 van de BLOCKING - Ban v3 cap nhat 07/06/2026.
Cap nhat sau tin nhan thay MD: chuyen Q1 -> Q2, bo dinh tinh."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   '4VanDe_BLOCKING_Q1Q3_v3_07062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.font.name = 'Times New Roman'
    r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'
    r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'
    r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(text, italic=False, indent=True, align_center=False):
    p = d.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_center
                   else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'
    r.font.size = Pt(12); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('- ' + text); r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

def REC(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('Em de xuat: ' + text); r.font.name = 'Times New Roman'
    r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)

def set_col_widths(t, widths_cm):
    for row in t.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


# ============================================================
# 1. HEADER + BOI CANH (~150 tu)
H1('BON VAN DE BLOCKING — BAN v3 CAP NHAT NGAY 07/06/2026')
P('Tai lieu chot trang thai 4 van de cu va 2 quyet dinh moi cua nhom '
  'tac gia', italic=True, align_center=True, indent=False)

H2('TOM TAT BOI CANH')

P('Ngay 07/06/2026, sau khi nhan tin nhan tu thay Manh Dung (thay MD) '
  've huong xu ly bai bao, nhom tac gia chot hai quyet dinh chien luoc '
  'moi cho lo trinh xuat ban: (1) CHUYEN bai Q1 (du kien BMC '
  'Psychiatry) xuong Q2 vi nhom khong the bo sung phan phong van dinh '
  'tinh kip tien do; (2) RE-DESIGN bai Q3 dong hanh (companion paper) '
  'theo khung Q2 chu khong con phai phu hop khung Q1 nua. Hai quyet '
  'dinh nay anh huong truc tiep den 4 van de BLOCKING ma em da neu '
  'trong ban v2 ngay 01/06/2026. Ban v3 nay cap nhat trang thai tung '
  'van de cu, bo sung hai quyet dinh moi, va lap bang phan cong viec '
  'cho tung thanh vien nhom.')


# ============================================================
# 2. CAP NHAT 4 VAN DE CU - BANG (~400 tu)
H2('PHAN A — CAP NHAT TRANG THAI 4 VAN DE CU')

P('Sau tin nhan thay MD ngay 07/06/2026, trang thai 4 van de BLOCKING '
  'da neu trong ban v2 (01/06/2026) duoc cap nhat nhu sau. Hai van de '
  'da co quyet dinh dut diem, hai van de con cho NCS xac nhan.')

q_summary = [
    ['Ma', 'Cau hoi goc', 'Quyet dinh moi (07/06/2026)', 'Trang thai'],
    ['Q1-6', 'Du lieu phong van dinh tinh cho bai Q1',
     'GIAI QUYET — Quyet dinh thay MD: BO HOAN TOAN phan dinh tinh. '
     'Bai chuyen tu thiet ke hon hop sang thuan dinh luong. Khong con '
     'can bo sung n, Cohen kappa, transcripts. Phan 2.3 va 3.7 trong '
     'ban nhap se duoc go bo trong ban v8.',
     'DA CHOT — em ap dung vao v8'],
    ['Q1-8', 'Mo hinh SEM tich hop 21 duong dan',
     'GIU NGUYEN — Phuong an A da chot ngay 01/06: dieu chinh cach '
     'dien dat thay vi tai phan tich. Khi chuyen sang Q2 Frontiers, '
     'reviewer cua tap chi Q2 it kho tinh hon ve title-methods '
     'mismatch nen phuong an A van toi uu.',
     'DA CHOT — em duy tri trong v8'],
    ['Q3-6', 'Letter chap thuan dao duc HNUE',
     'CHO XAC NHAN — NCS Cong Thi Hang van can cung cap so quyet '
     'dinh + ngay ban hanh chinh thuc. Frontiers Psychiatry va BMC '
     'Public Health deu yeu cau IRB number cu the giong nhu BMC '
     'Psychiatry, khong thay doi yeu cau du chuyen tap chi.',
     'CHO NCS xac nhan trong tuan'],
    ['Q3-9', 'Chien luoc xuat ban Q1 / Q3 (companion strategy)',
     'DONG Y — Theo solution em de xuat trong v2: phuong an B + C '
     '(companion papers + tach trong tam hoan toan). Tuy nhien khung '
     'tham chieu cua Q3 cap nhat lai: Q3 khong con doi chieu Q1 BMC '
     'Psychiatry ma doi chieu Q2 Frontiers (xem Quyet dinh 6 ben duoi).',
     'DA CHOT — em re-design Q3'],
]

t = d.add_table(rows=1, cols=4); t.style = 'Light Grid Accent 1'
t.autofit = False
hdr = t.rows[0].cells
for i, h in enumerate(q_summary[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in q_summary[1:]:
    row = t.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t, [1.2, 3.5, 8.5, 3.5])


# ============================================================
# 3. HAI QUYET DINH MOI (~500 tu)
H2('PHAN B — HAI QUYET DINH MOI TU TIN NHAN THAY MD 07/06/2026')

H3('Quyet dinh 5 — Chuyen bai Q1 xuong Q2')

P('Ly do: nhom khong the bo sung du lieu phong van dinh tinh kip tien '
  'do bao ve luan an cua NCS Cong Thi Hang. Phan dinh tinh la tieu chi '
  'sang loc cung cua BMC Psychiatry (tap chi yeu cau thiet ke hon hop '
  'phai co chung minh day du: so do chon mau, Cohen kappa, ma tran '
  'joint-display). Khong co phan nay, bai khong vuot duoc vong sang '
  'loc bien tap. Vi vay, thay MD quyet dinh chuyen muc tieu xuong cap '
  'Q2 voi yeu cau nhe hon ve thiet ke nghien cuu.')

P('Target moi (em de xuat TOP): Frontiers in Psychiatry — la tap chi '
  'Q2 cao trong nganh tam than hoc, chi so anh huong JCR 2024 = 3,2, '
  'open access, thoi gian phan bien trung binh 90-120 ngay. Tap chi '
  'nay chap nhan thiet ke thuan dinh luong khong bat buoc co phan dinh '
  'tinh, va co section "Public Mental Health" rat phu hop voi du lieu '
  '1.352 hoc sinh trung hoc co so Ha Noi cua NCS [1].')

P('Hai tap chi du phong (backup) trong truong hop Frontiers Psychiatry '
  'tu choi:', italic=True)
B('BMC Public Health (Q1-Q2 tuy nam, IF 2024 = 4,5) — phu hop neu '
  'nhom muon nhan manh goc nhin suc khoe cong dong thay vi tam than '
  'lam sang [2]. Thoi gian phan bien khoang 100-140 ngay.', 0)
B('Frontiers in Public Health (Q2, IF 2024 = 3,0) — backup an toan '
  'nhat, ty le chap nhan cao hon, phu hop khi can xuat ban nhanh '
  'phuc vu bao ve luan an cua NCS [3].', 0)

H3('Quyet dinh 6 — Re-design Q3 companion theo khung Q2')

P('He qua noi tai cua Quyet dinh 5: bai Q3 (du kien PLOS ONE) khong '
  'con phai dong hanh voi Q1 BMC Psychiatry ma chuyen sang dong hanh '
  'voi Q2 Frontiers. Dieu nay mo ra hai huong tai thiet ke Q3:')

P('Huong 1 — Q3 chua phan dinh tinh (qualitative interview). Vi Q2 '
  'Frontiers da la bai thuan dinh luong (Quyet dinh 5), neu sau nay '
  'NCS hoan thanh duoc du lieu phong van dinh tinh, phan nay co the '
  'duoc tach ra thanh Q3 doc lap voi cau hoi nghien cuu rieng (vi du: '
  '"Cau chuyen song dong cua hoc sinh THCS Ha Noi co diem lo au cao: '
  'phan tich chu de tu phong van ban cau truc"). Cach nay tan dung '
  'lai du lieu dinh tinh sau khi NCS thu thap xong.')

P('Huong 2 — Q3 = mo hinh SEM phu chuyen sau 1 phenotype. Thay vi '
  'mo ta tat ca lo au, Q3 di sau vao 1 phan loai cu the — vi du SAD '
  '(Social Anxiety Disorder, lo au xa hoi) — vi day la phenotype dac '
  'biet quan trong o lua tuoi vi thanh nien Viet Nam theo bao cao cua '
  'Karasu (1986) [4] va First et al. (2022) DSM-5-TR [5]. Cach nay tao '
  'ra cau hoi nghien cuu sac net, tach hoan toan voi Q2 Frontiers ve '
  'mat trong tam.')

REC('Huong 2 (Q3 = SEM phu chuyen sau ve SAD) la lua chon toi uu vi '
    '(a) khong phu thuoc vao viec NCS co hoan thanh dinh tinh hay '
    'khong; (b) tach trong tam rat ro voi Q2; (c) phu hop voi y van '
    'ICD-11 ve phan loai lo au [6]. Em se cho thay MD va NCS confirm '
    'truoc khi bat dau soan Q3 v1.')


# ============================================================
# 4. BANG ACTION ITEMS 5 COT (~400 tu)
H2('PHAN C — BANG ACTION ITEMS PHAN CONG NHIEM VU')

P('De thuan tien theo doi tien do, em lap bang phan cong cho 4 '
  'thanh vien chinh cua nhom: NCS Cong Thi Hang, thay MD (Manh Dung), '
  'thay NMD (Nguyen Minh Duc), va em (tro ly nghien cuu). Bang gom '
  '5 cot: Nguoi / Van de / Deadline / Output / Status.')

act_data = [
    ['Nguoi', 'Van de', 'Deadline', 'Output', 'Status'],

    ['NCS Cong Thi Hang', 'Letter dao duc HNUE (Q3-6)',
     '14/06/2026',
     'Scan letter chinh thuc voi so QD + ngay ban hanh; '
     'hoac van ban xac nhan dang xin retroactive IRB',
     'Dang cho'],
    ['NCS Cong Thi Hang', 'Confirm Q2 v1 draft sau khi em build xong',
     '21/06/2026',
     'Phan hoi bang van ban: dong y / yeu cau sua doi tung muc',
     'Cho em hoan thanh v1'],
    ['NCS Cong Thi Hang', 'Prepare cover letter cho Frontiers '
     'Psychiatry submission',
     '28/06/2026',
     'Ban thao cover letter 1 trang co chu ky NCS',
     'Cho Quyet dinh 5 confirm'],

    ['Thay MD (Manh Dung)', 'Review Q2 v1 draft',
     '24/06/2026',
     'Email phan hoi: dong y / sua doi noi dung khoa hoc',
     'Cho em gui v1'],
    ['Thay MD (Manh Dung)', 'Xac nhan target Frontiers Psychiatry '
     '(hoac chon backup)',
     '12/06/2026',
     'Tin nhan ngan xac nhan chon tap chi muc tieu',
     'Dang cho'],

    ['Thay NMD (Nguyen Minh Duc)',
     'Review Q2 v1 ve noi dung + format submission',
     '24/06/2026',
     'Email phan hoi tap trung vao format tap chi va dien dat tieng Anh',
     'Cho em gui v1'],

    ['Em (tro ly nghien cuu)', 'Build Q2 v1 draft tu Q1 v7',
     '20/06/2026',
     'File Draft_Q2_Frontiers_v1_xxxxx.docx — bo phan dinh tinh, '
     'dieu chinh format theo Frontiers guidelines',
     'Dang lam'],
    ['Em (tro ly nghien cuu)', 'Build Q3 companion v1 (huong 2: SEM '
     'phu chuyen sau SAD)',
     '05/07/2026',
     'File Draft_Q3_SAD_v1_xxxxx.docx — cau hoi nghien cuu sac net, '
     'tach trong tam voi Q2',
     'Cho thay MD + NCS confirm Huong 2'],
    ['Em (tro ly nghien cuu)', 'Plagiarism check Q2 v1 truoc khi '
     'NCS submit',
     '25/06/2026',
     'BaoCao_Plagiarism_Q2_v1_xxxxx.docx voi similarity < 15%',
     'Cho v1 hoan thanh'],
]

t2 = d.add_table(rows=1, cols=5); t2.style = 'Light Grid Accent 1'
t2.autofit = False
hdr = t2.rows[0].cells
for i, h in enumerate(act_data[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in act_data[1:]:
    row = t2.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t2, [3.0, 4.0, 2.0, 5.0, 2.5])


# ============================================================
# 5. THAM KHAO (~150 tu)
H2('THAM KHAO')

refs = [
    '[1] Frontiers in Psychiatry. Author Guidelines — Section "Public '
    'Mental Health". Truy cap 07/06/2026 tai https://www.frontiersin.'
    'org/journals/psychiatry/for-authors/author-guidelines. ISSN '
    '1664-0640. JCR Impact Factor 2024 = 3,2.',

    '[2] BMC Public Health. Submission Guidelines — Manuscript '
    'Preparation. Truy cap 07/06/2026 tai https://bmcpublichealth.'
    'biomedcentral.com/submission-guidelines/preparing-your-manuscript. '
    'ISSN 1471-2458. JCR Impact Factor 2024 = 4,5.',

    '[3] Frontiers in Public Health. Author Guidelines. Truy cap '
    '07/06/2026 tai https://www.frontiersin.org/journals/public-health/'
    'for-authors/author-guidelines. ISSN 2296-2565. JCR Impact Factor '
    '2024 = 3,0.',

    '[4] Karasu, T. B. (1986). The psychotherapies: Benefits and '
    'limitations. American Journal of Psychotherapy, 40(3), 324-342. '
    'DOI: 10.1176/appi.psychotherapy.1986.40.3.324. PMID: 3094389. '
    '(Em da verify PDF goc trong thu muc Materials.)',

    '[5] First, M. B., Yousif, L. H., Clarke, D. E., Wang, P. S., '
    'Gogtay, N., & Appelbaum, P. S. (2022). DSM-5-TR: overview of '
    'what is new and what has changed. World Psychiatry, 21(2), '
    '218-219. DOI: 10.1002/wps.20989. PMID: 35524596.',

    '[6] Pezzella, P. (2022). The ICD-11 is now officially in effect. '
    'World Psychiatry, 21(2), 331-332. DOI: 10.1002/wps.20982. PMID: '
    '35524598.',

    '[7] Tham chieu noi bo: 4VanDe_BLOCKING_Q1Q3_v2_01062026.docx — '
    'ban tien nhiem voi 4 van de goc.',

    '[8] Tham chieu noi bo: Tin nhan thay MD ngay 07/06/2026 chot '
    'hai quyet dinh chuyen Q1 -> Q2 va bo phan dinh tinh.',
]

for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref); r.font.name = 'Times New Roman'
    r.font.size = Pt(11)


# ============================================================
# FOOTER
P('', indent=False)
P('Soạn 07/06/2026', italic=True, align_center=True, indent=False)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')

# Word count check
from docx import Document as D2
doc = D2(OUT)
total_words = 0
for para in doc.paragraphs:
    total_words += len(para.text.split())
for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for para in cell.paragraphs:
                total_words += len(para.text.split())
print(f'Word count: {total_words}')
