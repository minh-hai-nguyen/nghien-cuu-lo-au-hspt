"""Build doc binh luan ve muc do OR (parenting + resilience)."""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/OR_muc_do_binh_luan_parenting_resilience.docx')

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x66, 0x66, 0x66)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    """Đoạn tô xanh — dùng cho câu hỏi và CÂU TRẢ LỜI."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = BLUE
    r.font.size = Pt(12)
    r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    """Đoạn đen — dùng cho background."""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    r.font.size = Pt(12)
    r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.color.rgb = BLUE
    r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.italic = italic
    return p

# =================================================================
# TIÊU ĐỀ
# =================================================================
H('Bình luận về mức độ Odds Ratio (OR) — Parenting & Resilience', level=1)
H('(Nuôi dạy con + khả năng phục hồi & lo âu/trầm cảm)', level=2)

# =================================================================
# CÂU HỎI (tô xanh)
# =================================================================
H('Câu hỏi của thầy', level=2)
para_blue('Em bình luận cho thầy câu này nhé. Mức độ có cao không hay là thấp?')
para_blue('"Mô hình \'nuôi dạy tiêu cực\' có mối liên quan thuận với triệu chứng lo âu (OR = 2,01; KTC 95%: 1,38–2,92). Ngược lại, mô hình \'nuôi dạy tích cực\' có mối liên quan nghịch với triệu chứng lo âu (OR = 0,32; KTC 95%: 0,24–0,43)."')
para_blue('"Mức khả năng phục hồi thấp có mối liên quan thuận với trầm cảm (OR = 6,74; KTC 95%: 4,73–9,61) và triệu chứng lo âu (OR = 2,80; KTC 95%: 1,92–4,09)."')

# =================================================================
# BACKGROUND (đen) — đặt nền tảng để hiểu trả lời
# =================================================================
H('1. Hiểu nhanh về OR (Odds Ratio)', level=2)
para_black('OR là tỷ số chênh — đo lường mức độ liên quan giữa một yếu tố (exposure) và một kết cục (outcome). Cách diễn giải nhanh:')
bullet_black('OR = 1: KHÔNG có mối liên quan.')
bullet_black('OR > 1: yếu tố LÀM TĂNG nguy cơ kết cục. OR = 2 nghĩa là người có yếu tố đó có nguy cơ (theo odds) cao GẤP 2 LẦN người không có.')
bullet_black('OR < 1: yếu tố LÀM GIẢM nguy cơ — đây là yếu tố BẢO VỆ. OR = 0,32 nghĩa là người có yếu tố đó chỉ có 32% odds so với nhóm không có; tương đương GIẢM 68% odds.')
bullet_black('Khi OR < 1, để dễ so sánh "độ mạnh" với OR > 1, lấy nghịch đảo: 1/0,32 ≈ 3,13.')

H('2. Bảng phân loại "độ mạnh" của OR (đồng thuận chung trong dịch tễ học)', level=2)
para_black('Trong các nghiên cứu epidemiology/public health về tâm lý-tâm thần, mức độ thường được phân nhóm như sau (Cohen, Hosmer-Lemeshow, Rosenthal):')

tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Light Grid Accent 1'
header = tbl.rows[0]
header.cells[0].text = 'OR (hoặc 1/OR nếu < 1)'
header.cells[1].text = 'Mức độ'
header.cells[2].text = 'Ý nghĩa thực tế'
data = [
    ('1,0 – 1,5', 'Rất yếu (negligible)', 'Có thể do nhiễu hoặc chance, không có ý nghĩa thực tế nhiều'),
    ('1,5 – 2,5', 'Yếu (weak / small)', 'Có liên quan nhưng nhỏ; cần kiểm soát thêm yếu tố nhiễu'),
    ('2,5 – 4,0', 'Trung bình (moderate)', 'Liên quan rõ rệt, có ý nghĩa thực hành'),
    ('4,0 – 6,0', 'Mạnh (strong)', 'Liên quan đáng kể, cần can thiệp'),
    ('> 6,0', 'Rất mạnh (very strong)', 'Mức độ liên quan cao; thường gặp trong các yếu tố nguy cơ lớn'),
]
for i, (a, b, c) in enumerate(data, 1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black('Lưu ý: bảng này là CHỈ DẪN THAM KHẢO, không tuyệt đối. Trong tâm lý học, OR ~ 2 đã được nhiều nghiên cứu coi là "đáng quan tâm". Các so sánh trên nguy cơ tuyệt đối (absolute risk) và sự biến đổi theo bối cảnh quần thể cũng quan trọng.', italic=True)

# =================================================================
# PHÂN TÍCH 4 OR — vẫn phần background đen
# =================================================================
H('3. Phân tích từng OR trong câu của thầy', level=2)

H('3.1. Nuôi dạy tiêu cực ↔ Triệu chứng lo âu — OR = 2,01 (KTC 95%: 1,38–2,92)', level=3)
para_black('• OR = 2,01: Trẻ có cha mẹ áp dụng phong cách "nuôi dạy tiêu cực" (kiểm soát quá mức, từ chối, trừng phạt khắc nghiệt) có odds bị triệu chứng lo âu CAO GẤP 2,01 LẦN trẻ không có yếu tố này.')
para_black('• KTC 95% từ 1,38 đến 2,92: KHÔNG chứa giá trị 1 → có ý nghĩa thống kê (p < 0,05). Khoảng tin cậy KHÁ RỘNG (1,38–2,92), gợi ý cỡ mẫu chưa đủ lớn hoặc có dao động cao.')
para_black('• Mức độ: rơi vào nhóm YẾU – TRUNG BÌNH (1,5–2,5).')

H('3.2. Nuôi dạy tích cực ↔ Triệu chứng lo âu — OR = 0,32 (KTC 95%: 0,24–0,43)', level=3)
para_black('• OR < 1 → đây là YẾU TỐ BẢO VỆ (protective factor).')
para_black('• OR = 0,32: Trẻ có cha mẹ áp dụng phong cách "nuôi dạy tích cực" (ấm áp, hỗ trợ, đồng hành) chỉ có odds bị lo âu bằng 32% so với trẻ không có yếu tố này → giảm 68% odds.')
para_black('• Lấy nghịch đảo để so sánh độ mạnh: 1/0,32 ≈ 3,13 → mức TRUNG BÌNH (2,5–4,0).')
para_black('• KTC 95% từ 0,24 đến 0,43: chặt chẽ, không chứa 1 → kết quả có ý nghĩa thống kê và đáng tin cậy. Tương đương 1/0,43 ≈ 2,33 đến 1/0,24 ≈ 4,17 — khoảng vẫn nằm trong YẾU đến MẠNH.')
para_black('• So với OR = 2,01 ở phía "tiêu cực", thì protective effect của "tích cực" (1/0,32 ≈ 3,13) MẠNH HƠN khoảng 1,5 lần. Đây là phát hiện quan trọng: NUÔI DẠY TÍCH CỰC bảo vệ mạnh hơn nuôi dạy tiêu cực gây hại.')

H('3.3. Khả năng phục hồi thấp ↔ TRẦM CẢM — OR = 6,74 (KTC 95%: 4,73–9,61)', level=3)
para_black('• OR = 6,74: trẻ có khả năng phục hồi (resilience) thấp có odds bị TRẦM CẢM CAO GẤP 6,74 LẦN trẻ phục hồi cao.')
para_black('• Đây là OR RẤT CAO — vượt ngưỡng 6,0 → mức RẤT MẠNH (very strong).')
para_black('• KTC 95% từ 4,73 đến 9,61: cả khoảng đều nằm trong vùng "mạnh" đến "rất mạnh"; không chứa 1 → có ý nghĩa thống kê chắc chắn.')
para_black('• Đây là FINDING NỔI BẬT NHẤT trong câu của thầy. Trong tâm lý học vị thành niên, OR ~ 7 đối với trầm cảm là mức RẤT đáng chú ý — gợi ý resilience là yếu tố can thiệp ưu tiên.')

H('3.4. Khả năng phục hồi thấp ↔ Triệu chứng LO ÂU — OR = 2,80 (KTC 95%: 1,92–4,09)', level=3)
para_black('• OR = 2,80: trẻ phục hồi thấp có odds bị lo âu cao gấp 2,8 lần trẻ phục hồi cao.')
para_black('• Mức TRUNG BÌNH (2,5–4,0) — vẫn rõ rệt và có ý nghĩa lâm sàng.')
para_black('• KTC 95% từ 1,92 đến 4,09: kéo dài từ YẾU đến MẠNH; khá rộng nhưng không chứa 1.')
para_black('• Quan sát thú vị: cùng yếu tố "phục hồi thấp" nhưng tác động lên TRẦM CẢM (OR=6,74) MẠNH HƠN GẤP HƠN 2 LẦN tác động lên LO ÂU (OR=2,80). Điều này phù hợp với y văn — resilience đặc biệt quan trọng để phòng trầm cảm; lo âu chịu ảnh hưởng nhiều bởi các yếu tố nhận thức-sinh học khác.')

# =================================================================
# CÂU TRẢ LỜI (tô xanh, gom trước phụ lục)
# =================================================================
H('4. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Tóm gọn cho thầy: trong 4 OR, có 2 OR mức TRUNG BÌNH, 1 OR mức YẾU – TRUNG BÌNH, và 1 OR mức RẤT MẠNH. Cụ thể:', bold=True)

bullet_blue('Nuôi dạy TIÊU CỰC ↔ lo âu: OR = 2,01 → mức YẾU – TRUNG BÌNH. Có ý nghĩa thống kê nhưng không quá lớn; tăng khoảng gấp đôi nguy cơ.')

bullet_blue('Nuôi dạy TÍCH CỰC ↔ lo âu: OR = 0,32 (= 1/3,13) → mức TRUNG BÌNH (protective). Hiệu ứng bảo vệ MẠNH HƠN tác động gây hại của nuôi dạy tiêu cực — đây là phát hiện đáng nhấn mạnh.')

bullet_blue('Phục hồi thấp ↔ TRẦM CẢM: OR = 6,74 → mức RẤT MẠNH. Đây là OR cao nhất trong cả nhóm; cho thấy khả năng phục hồi (resilience) là yếu tố BẢO VỆ KEY chống trầm cảm vị thành niên.')

bullet_blue('Phục hồi thấp ↔ LO ÂU: OR = 2,80 → mức TRUNG BÌNH. Vẫn rõ rệt nhưng không "đậm" như đối với trầm cảm.')

para_blue('Tổng kết về mức độ:', bold=True)
para_blue('Hầu hết các OR ở câu này nằm trong khoảng YẾU đến TRUNG BÌNH (OR 2–3), tức là KHÔNG QUÁ CAO so với mặt bằng nghiên cứu dịch tễ tâm lý. Riêng OR = 6,74 (resilience thấp ↔ trầm cảm) là RẤT CAO và là phát hiện nổi bật nhất, đáng đưa vào báo cáo can thiệp như "ưu tiên nâng cao khả năng phục hồi".')

para_blue('Hàm ý ứng dụng cho can thiệp:', bold=True)
bullet_blue('Vì OR (resilience↔depression) lên tới 6,74, các chương trình tăng cường resilience (CBT, mindfulness, đào tạo kỹ năng đối phó) có thể có IMPACT LỚN trên trầm cảm hơn là can thiệp lên parenting.')
bullet_blue('Vì protective effect của nuôi dạy tích cực (1/0,32 ≈ 3,13) MẠNH HƠN tác động xấu của nuôi dạy tiêu cực (OR = 2,01), các chương trình tập huấn phụ huynh nên tập trung XÂY DỰNG NUÔI DẠY TÍCH CỰC, không chỉ là "tránh nuôi dạy tiêu cực".')
bullet_blue('Tất cả KTC 95% đều không chứa 1 → các phát hiện đều có ý nghĩa thống kê đáng tin cậy.')

para_blue('Một lưu ý cẩn trọng:', bold=True)
para_blue('OR là tương đối, không phải tuyệt đối. Khi tỷ lệ kết cục trong quần thể CAO (>10%, ví dụ lo âu ở HS THCS/THPT), OR sẽ phóng đại so với risk ratio (RR) thật. Khi áp dụng vào tư vấn, nên nói "tăng/giảm odds gấp X lần" thay vì "tăng/giảm nguy cơ gấp X lần" để chính xác. Cần tham khảo prevalence của outcome trong dân số gốc của nghiên cứu mới biết absolute risk thực tế là bao nhiêu.')

# =================================================================
# PHỤ LỤC — REFERENCES (cuối doc)
# =================================================================
H('5. Phụ lục — Tài liệu tham khảo', level=2)

para_black('1. Cohen, J. (1988). Statistical Power Analysis for the Behavioral Sciences (2nd ed.). Hillsdale, NJ: Lawrence Erlbaum Associates. — Khung phân loại effect size cổ điển; OR khoảng 1,5/3,5/9 được Chen et al. (2010) ánh xạ tương ứng với d = 0,2/0,5/0,8 (small/medium/large).', italic=True)

para_black('2. Chen, H., Cohen, P., & Chen, S. (2010). How big is a big odds ratio? Interpreting the magnitudes of odds ratios in epidemiological studies. Communications in Statistics – Simulation and Computation, 39(4), 860–864. doi:10.1080/03610911003650383 — Bài tham chiếu cốt lõi cho việc dịch OR sang Cohen d và phân loại độ mạnh.', italic=True)

para_black('3. Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). Applied Logistic Regression (3rd ed.). Hoboken, NJ: Wiley. — Sách tham khảo chuẩn về logistic regression và diễn giải OR.', italic=True)

para_black('4. Rosenthal, J. A. (1996). Qualitative descriptors of strength of association and effect size. Journal of Social Service Research, 21(4), 37–59. — Đề xuất phân loại "negligible/weak/moderate/strong" được dùng rộng rãi trong xã hội và sức khỏe.', italic=True)

para_black('5. Tham khảo trong corpus dự án (DB Lo-au):')
bullet_black('VN013 Tô Thị Hồng (2017) — DASS-21 trên HS THCS Hà Nội (về parenting và lo âu).', italic=True)
bullet_black('VN016 Nguyễn Cao Minh (2012) — RCADS chuẩn hóa VN; có tỷ lệ prevalence để tính absolute risk.', italic=True)
bullet_black('QT059 (mindfulness/resilience meta-analysis trong DB) — pooled effect size về resilience.', italic=True)
bullet_black('QT067 Pascoe et al. (2020) — narrative review nhấn mạnh protective effect của resilience trên depression.', italic=True)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
