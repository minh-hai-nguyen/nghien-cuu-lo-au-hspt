"""Build doc tra loi: ty le HS chu dong tim gap bac si tam ly o truong."""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Ty_le_HS_tim_gap_bac_si_tam_ly_o_truong.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(12); r.italic = italic
    return p

# =========================================================
# TIÊU ĐỀ
# =========================================================
H('Tỷ lệ HS chủ động tìm gặp bác sĩ tâm lý ở trường', level=1)
H('— so với tổng số HS có vấn đề lo âu/SKTT —', level=2)

# =========================================================
# CÂU HỎI (xanh)
# =========================================================
H('Câu hỏi của thầy', level=2)
para_blue(
    'Có tài liệu nào ước tính về tỷ lệ HS chủ động tìm gặp bác sĩ tâm lý ở trường so với '
    'tổng có vấn đề lo âu hoặc SKTT không?'
)

# =========================================================
# BACKGROUND (đen)
# =========================================================
H('1. Tổng quan: khái niệm "treatment gap" và "help-seeking"', level=2)

para_black(
    'Trong dịch tễ tâm thần, có hai khái niệm liên quan thầy cần phân biệt:'
)
bullet_black(
    'Treatment gap (khoảng cách điều trị) = % người có rối loạn nhưng KHÔNG nhận được '
    'bất kỳ dịch vụ chăm sóc nào.'
)
bullet_black(
    'Help-seeking (hành vi tìm kiếm sự giúp đỡ) = chia nhỏ theo loại nguồn hỗ trợ: '
    'chuyên gia tâm thần, bác sĩ đa khoa, nhân viên trường, gia đình, bạn bè, hoặc tự '
    'giải quyết.'
)
para_black(
    'Câu hỏi của thầy thuộc về help-seeking — cụ thể là "tìm gặp bác sĩ/chuyên gia tâm lý '
    'tại TRƯỜNG" (school-based mental health service). Đây là một subset rất nhỏ trong '
    'tổng help-seeking.'
)

H('2. Nguồn dữ liệu CHÍNH cho Việt Nam: VNAMHS 2022 (VN002)', level=2)

para_black(
    'V-NAMHS 2022 (Viet Nam Adolescent Mental Health Survey) là báo cáo điều tra QUỐC GIA '
    'ĐẦU TIÊN của VN về rối loạn SKTT vị thành niên 10–17 tuổi theo chuẩn DSM-5. UNICEF '
    'Việt Nam phối hợp Bộ LĐ-TB-XH + Tổng cục Thống kê + Viện Xã hội học. Mẫu đại diện '
    'quốc gia 5.996 cặp phụ huynh–vị thành niên từ 38 tỉnh/thành, 4 vùng địa lý. Đây là '
    'NGUỒN ĐÁNG TIN CẬY NHẤT cho VN hiện nay.'
)

H('Số liệu cốt lõi từ VN002 (mục "Sử dụng dịch vụ"):', level=3)

tbl = doc.add_table(rows=8, cols=2)
tbl.style = 'Light Grid Accent 1'
header = tbl.rows[0]
header.cells[0].text = 'Chỉ số'
header.cells[1].text = 'Tỷ lệ (weighted theo dân số VTN VN)'
data = [
    ('Tỷ lệ VTN có vấn đề SKTT (mental health problem, 12 tháng qua)', '21,7%'),
    ('Tỷ lệ VTN có rối loạn DSM-5 (full criteria)', 'Lo âu cao nhất; ~5–8% (clinical level)'),
    ('Trong nhóm có vấn đề SKTT — đã sử dụng BẤT KỲ dịch vụ nào', '8,4% (n = 389)'),
    ('Trong nhóm đã sử dụng — đến bác sĩ/điều dưỡng (cơ sở y tế)', '56,2%'),
    ('Trong nhóm đã sử dụng — đến nhân viên y tế cộng đồng', '10,7%'),
    ('Trong nhóm đã sử dụng — đến NHÂN VIÊN TRƯỜNG (school staff)', '5,5%'),
    ('Trong nhóm đã sử dụng — đến chuyên gia tâm thần (psychiatrist/psychologist)', '1,4%'),
]
for i, (a, b) in enumerate(data, 1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black('Một số phát hiện bổ sung từ VN002:', bold=True)
bullet_black('Chỉ 5,1% phụ huynh xác định được con cần trợ giúp; trong số đó chỉ 37,7% nhận đủ trợ giúp.')
bullet_black('Hỗ trợ phi chính thức (gia đình, bạn bè, tự giải quyết): 73,9% VTN.')
bullet_black('Rào cản hàng đầu: "tự giải quyết trong gia đình" 20,4%; "không chắc con có cần" 10,7%; "không biết chỗ tìm" 10,0%.')

H('3. Tính toán cụ thể cho câu hỏi của thầy', level=2)

para_black(
    'Câu hỏi: "Tỷ lệ HS chủ động tìm gặp bác sĩ tâm lý Ở TRƯỜNG / tổng số HS có vấn đề '
    'SKTT" = ?'
)

para_black('Hai bước nhân:', bold=True)
bullet_black('Bước 1: 8,4% trong số HS có vấn đề SKTT đã sử dụng BẤT KỲ dịch vụ nào.')
bullet_black('Bước 2: trong nhóm đã sử dụng, chỉ 5,5% đến nhân viên trường.')
para_black('=> 8,4% × 5,5% ≈ 0,46% — tức trong 100 HS có vấn đề SKTT, CHƯA TỚI 1 EM tìm gặp nhân viên/phòng tư vấn trường.', bold=True)

para_black('')
para_black(
    'Lưu ý: VN002 dùng cụm "school staff" (nhân viên trường) bao gồm cả giáo viên chủ nhiệm, '
    'nhân viên y tế trường, và tham vấn viên — KHÔNG phân tách riêng "bác sĩ tâm lý ở '
    'trường". Vì hệ thống tham vấn học đường ở VN còn rất thưa thớt, con số 5,5% chủ yếu '
    'là HS gặp giáo viên/y tá trường, KHÔNG hẳn là chuyên gia tâm lý chuyên nghiệp.'
)

H('4. Treatment gap tổng cho VN', level=2)
para_black(
    '100% − 8,4% = 91,6% HS có vấn đề SKTT KHÔNG nhận BẤT KỲ dịch vụ nào. Đây là '
    'treatment gap rất lớn, nằm ở mức cao trong khu vực Đông Nam Á.'
)

H('5. Các tài liệu BỔ SUNG trong corpus dự án', level=2)

bullet_black(
    'VN022 — UNICEF (2022) "Tâm điểm về Sức khỏe Tâm thần của Trẻ em Vị thành niên Nam '
    'và Nữ tại Việt Nam": khảo sát 668 HS THCS/THPT 5 tỉnh + 66 GV, có thông tin về '
    'stigma và rào cản help-seeking. Tham khảo bổ sung cho VN002.', italic=True
)
bullet_black(
    'QT008 — Wang et al. "A Latent Profile Analysis of Anxiety among Junior High School '
    'Students in Rural China": khuyến nghị "thiết lập phòng tư vấn tâm lý chuyên nghiệp '
    'tại trường" — gợi ý gap dịch vụ ở TQ cũng tương tự VN.', italic=True
)
bullet_black(
    'VN028 — Đào Thị Ngoãn và cs. (2025) ĐH Y Hà Nội: dùng "Bộ câu hỏi nhu cầu tham vấn '
    'tâm lý" trên SV năm 4 — có dữ liệu nhu cầu ở SV (không phải HS THCS/THPT).',
    italic=True
)
bullet_black(
    'QT066 — Murphy et al. (2024) Scoping review về peer support trong primary youth MH '
    'care (Ireland) — gián tiếp giải thích vì sao giới trẻ thường chọn bạn bè thay vì '
    'chuyên gia.', italic=True
)
bullet_black(
    'QT067 — Pascoe, Hetrick & Parker (2020): nhắc đến school connectedness ↔ less '
    'anxiety nhưng KHÔNG có % help-seeking cụ thể.', italic=True
)

H('6. Tài liệu QUỐC TẾ NGOÀI corpus (đáng cân nhắc bổ sung)', level=2)

bullet_black(
    'Gulliver, A., Griffiths, K. M., & Christensen, H. (2010). Perceived barriers and '
    'facilitators to mental health help-seeking in young people: A systematic review. BMC '
    'Psychiatry, 10, 113. — Tổng quan KINH ĐIỂN; rào cản hàng đầu: stigma, ngại bộc lộ, '
    'thích tự giải quyết.', italic=True
)
bullet_black(
    'Kessler, R. C., et al. (2009). The global burden of mental disorders: An update from '
    'the WHO World Mental Health (WMH) surveys. Epidemiology and Psychiatric Sciences. — '
    'Treatment gap toàn cầu ~50–80% ở thanh thiếu niên.', italic=True
)
bullet_black(
    'WHO (2021). Mental health atlas 2020. Geneva. — Số liệu so sánh quốc tế về service '
    'utilization.', italic=True
)
bullet_black(
    'Radez, J., et al. (2021). Why do children and adolescents (not) seek and access '
    'professional help for their mental health problems? A systematic review of '
    'quantitative and qualitative studies. European Child & Adolescent Psychiatry, 30, '
    '183–211. — Tổng quan mới; lý do trẻ TÌM hoặc KHÔNG tìm trợ giúp chuyên môn.',
    italic=True
)
bullet_black(
    'Health Behaviour in School-aged Children (HBSC) study — báo cáo định kỳ Châu Âu, có '
    'dữ liệu so sánh % HS gặp school nurse/counselor giữa các nước.', italic=True
)

# =========================================================
# CÂU TRẢ LỜI (xanh, gom 1 chỗ trước phụ lục)
# =========================================================
H('7. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Trả lời ngắn cho thầy:', bold=True)
para_blue(
    'CÓ — tài liệu chính cho VN là V-NAMHS 2022 (VN002 trong DB dự án), điều tra QUỐC GIA '
    'đầu tiên về SKTT vị thành niên VN. Đây là nguồn đáng tin cậy nhất hiện nay.'
)

para_blue('Số liệu cốt lõi (VN002, weighted, đại diện toàn quốc VN):', bold=True)
bullet_blue('21,7% VTN VN có vấn đề SKTT trong 12 tháng qua.')
bullet_blue('Trong nhóm này, chỉ 8,4% đã sử dụng bất kỳ dịch vụ nào → treatment gap = 91,6%.')
bullet_blue('Trong nhóm đã sử dụng dịch vụ, chỉ 5,5% đến NHÂN VIÊN TRƯỜNG.')
bullet_blue('Trong nhóm đã sử dụng dịch vụ, chỉ 1,4% đến CHUYÊN GIA TÂM THẦN.')

para_blue('Tính toán trực tiếp cho câu hỏi của thầy:', bold=True)
bullet_blue(
    'Tỷ lệ HS có vấn đề SKTT mà CHỦ ĐỘNG TÌM tới nhân viên/phòng tư vấn TRƯỜNG '
    '= 8,4% × 5,5% ≈ 0,46% — tức TRONG 100 HS CÓ VẤN ĐỀ SKTT, CHƯA TỚI 1 EM tìm tới '
    'trường để được hỗ trợ.'
)
bullet_blue(
    'Nếu tính chuyên gia tâm thần riêng: 8,4% × 1,4% ≈ 0,12% — tức 12/10.000 HS có vấn '
    'đề SKTT mới gặp được chuyên gia thực thụ.'
)

para_blue('Cảnh báo về cách hiểu:', bold=True)
bullet_blue(
    '"School staff" trong VN002 bao gồm cả giáo viên chủ nhiệm, y tá trường, tham vấn '
    'viên — KHÔNG phân tách riêng "bác sĩ tâm lý chuyên nghiệp ở trường". Vì hệ thống '
    'tham vấn học đường VN còn thưa, con số 5,5% chủ yếu là HS gặp GV/y tá, không hẳn '
    'chuyên gia tâm lý.'
)
bullet_blue(
    'Để có số liệu CHÍNH XÁC HƠN về "bác sĩ tâm lý ở trường" riêng biệt, cần khảo sát '
    'bổ sung — VN hiện chưa có bài nào tách rõ. Có thể dựa vào VN022 (UNICEF 2022) cho '
    'data khu vực, hoặc Gulliver 2010 / Radez 2021 cho khung phân loại quốc tế.'
)

para_blue('Hàm ý cho đề tài lo âu HS THCS/THPT của thầy:', bold=True)
bullet_blue(
    'Treatment gap 91,6% là CON SỐ LỚN, đáng đưa vào phần "rationale" của đề cương — '
    'minh họa tính cấp thiết của can thiệp dựa vào trường (school-based intervention).'
)
bullet_blue(
    'Nếu phát triển chương trình tư vấn học đường, baseline ~0,5% HS chủ động tìm gặp '
    'có thể là KPI dễ tăng — chỉ cần đạt 5–10% là đã cải thiện 10–20 lần.'
)
bullet_blue(
    'Khi viết: nên dẫn VN002 trực tiếp + so sánh với GSHS 2019 (ý nghĩ tự sát 15,6% — '
    'cao hơn VN002 do GSHS dùng tự khai vs VN002 phỏng vấn face-to-face).'
)

para_blue('Tóm lại:', bold=True)
para_blue(
    'CÓ tài liệu (chính là VN002 VNAMHS 2022). Tỷ lệ HS chủ động tìm gặp tới nhân viên '
    'trường ≈ 0,46%. Treatment gap tổng = 91,6%. Đây là phát hiện rất quan trọng cho đề '
    'cương can thiệp.'
)

# =========================================================
# PHỤ LỤC — REFERENCES
# =========================================================
H('8. Phụ lục — Tài liệu tham khảo', level=2)

para_black(
    '1. Viện Xã hội học, UNICEF Việt Nam, Bộ LĐ-TB-XH, Tổng cục Thống kê (2022). Khảo '
    'sát Sức khỏe Tâm thần Vị thành niên Việt Nam (V-NAMHS 2022). Hà Nội. — '
    '(VN002 trong DB dự án) Báo cáo điều tra quốc gia đầu tiên về SKTT VTN VN; n = '
    '5.996 cặp phụ huynh-vị thành niên; chuẩn DSM-5 (DISC-5).', italic=True
)

para_black(
    '2. UNICEF Việt Nam (2022). Tâm điểm về Sức khỏe Tâm thần của Trẻ em Vị thành niên '
    'Nam và Nữ tại Việt Nam (Adolescent Boys and Girls in Viet Nam). Hà Nội. — (VN022 '
    'trong DB dự án) Khảo sát 668 HS THCS/THPT + 66 GV ở 5 tỉnh; mixed methods.',
    italic=True
)

para_black(
    '3. Gulliver, A., Griffiths, K. M., & Christensen, H. (2010). Perceived barriers and '
    'facilitators to mental health help-seeking in young people: A systematic review. '
    'BMC Psychiatry, 10, 113. doi:10.1186/1471-244X-10-113 — Tổng quan kinh điển về rào '
    'cản và yếu tố thúc đẩy help-seeking ở thanh thiếu niên.', italic=True
)

para_black(
    '4. Radez, J., Reardon, T., Creswell, C., Lawrence, P. J., Evdoka-Burton, G., & '
    'Waite, P. (2021). Why do children and adolescents (not) seek and access '
    'professional help for their mental health problems? A systematic review of '
    'quantitative and qualitative studies. European Child & Adolescent Psychiatry, '
    '30(2), 183–211. doi:10.1007/s00787-019-01469-4 — Tổng quan cập nhật về help-seeking.',
    italic=True
)

para_black(
    '5. Kessler, R. C., Aguilar-Gaxiola, S., Alonso, J., Chatterji, S., Lee, S., Ormel, '
    'J., Üstün, T. B., & Wang, P. S. (2009). The global burden of mental disorders: An '
    'update from the WHO World Mental Health (WMH) surveys. Epidemiology and '
    'Psychiatric Sciences, 18(1), 23–33. — Treatment gap toàn cầu.',
    italic=True
)

para_black(
    '6. World Health Organization (2021). Mental health atlas 2020. Geneva: WHO. — Dữ '
    'liệu service utilization quốc tế.', italic=True
)

para_black(
    '7. Wang, X., et al. (2022/2023). A Latent Profile Analysis of Anxiety among Junior '
    'High School Students in Rural China. International Journal of Environmental Research '
    'and Public Health. — (QT008 trong DB dự án) Đề xuất phòng tư vấn tâm lý ở trường.',
    italic=True
)

para_black(
    '8. Đào Thị Ngoãn và cs. (2025). Thực trạng tâm lý của sinh viên năm thứ tư Trường '
    'Đại học Y Hà Nội năm học 2023-2024. — (VN028 trong DB dự án) Có "Bộ câu hỏi nhu '
    'cầu tham vấn tâm lý" — dữ liệu nhu cầu ở SV.', italic=True
)

para_black(
    '9. Murphy, R., et al. (2024). Peer support in primary youth mental health care: '
    'A scoping review. Journal of Community Psychology, 52(1), 154–180. doi:10.1002/'
    'jcop.23090 — (QT066 trong DB dự án) Lý do giới trẻ chọn peer thay vì chuyên gia.',
    italic=True
)

para_black(
    '10. Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on '
    'students in secondary school and higher education. International Journal of '
    'Adolescence and Youth, 25(1), 104–112. doi:10.1080/02673843.2019.1596823 — '
    '(QT067 trong DB dự án) Có nhắc school connectedness; không có % cụ thể.', italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
