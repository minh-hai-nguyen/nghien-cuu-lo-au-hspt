# -*- coding: utf-8 -*-
"""Doc 4 van de BLOCKING - gui thay HD + NCS Cong Thi Hang.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', '4VanDe_BLOCKING_Q1Q3_v2_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(text, italic=False, indent=True, align_center=False):
    p = d.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_center
                   else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('▸ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

def WARN(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def REC(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('★ Em đề xuất: ' + text); r.font.name = 'Times New Roman'
    r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)

def set_col_widths(t, widths_cm):
    for row in t.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


# ============================================================
H1('BỐN VẤN ĐỀ — TRẠNG THÁI SAU KHI NHÓM QUYẾT ĐỊNH')
P('Bản v2 cập nhật quyết định của nhóm tác giả ngày 01/06/2026',
  italic=True, align_center=True, indent=False)
P('Tham chiếu cho bản nháp bài Q1 (BMC Psychiatry / Q2 contingency) '
  'và bài Q3 (PLOS ONE)', italic=True, align_center=True, indent=False)


# ============================================================
H2('TÓM TẮT TRẠNG THÁI BỐN VẤN ĐỀ')

q_summary = [
    ['Mã', 'Câu hỏi', 'Quyết định của nhóm', 'Trạng thái'],
    ['Q1-6', 'Dữ liệu phỏng vấn định tính cho bài Q1', 'Khó bổ sung '
     'theo yêu cầu Q1 → cân nhắc CHUYỂN BÀI Q1 XUỐNG Q2 nếu phần '
     'định tính không kịp', '⚠ Chờ NCS đánh giá khả năng bổ sung'],
    ['Q1-8', 'Mô hình SEM tích hợp 21 đường dẫn', 'PHƯƠNG ÁN A — '
     'giữ nguyên phân tích, điều chỉnh cách diễn đạt (không tái '
     'phân tích)', '✓ Đã quyết, em đã áp dụng vào Draft Q1 v3'],
    ['Q3-6', 'Letter chấp thuận đạo đức HNUE', 'Chưa rõ', '⚠ Chờ NCS '
     'cung cấp số quyết định + ngày'],
    ['Q3-9', 'Chiến lược nộp bài Q1 ↔ Q3', 'PHƯƠNG ÁN B + C — em '
     'viết Q3 song song với Q1/Q2, công bố đồng thời (companion '
     'papers) + tách trọng tâm hoàn toàn', '✓ Đã quyết, em bắt '
     'đầu soạn Draft Q3 sau khi xong Q1 v4'],
]

t = d.add_table(rows=1, cols=4); t.style = 'Light Grid Accent 1'
t.autofit = False
hdr = t.rows[0].cells
for i, h in enumerate(q_summary[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in q_summary[1:]:
    row = t.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
def _set_col_widths(t, widths_cm):
    for row in t.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)
_set_col_widths(t, [1.2, 5.0, 7.0, 3.5])


H2('LỜI MỞ ĐẦU')

P('Tài liệu này được soạn theo cách cho phép người đọc bắt đầu mà '
  'không cần đọc trước bản nháp bài báo. Mỗi vấn đề được trình bày '
  'theo cấu trúc thống nhất gồm năm phần: (a) "Bối cảnh" giúp hiểu '
  'tình huống hiện tại; (b) "Tại sao quan trọng" giải thích hệ quả '
  'nếu không quyết định; (c) "Câu hỏi cụ thể cho NCS/thầy"; (d) "Các '
  'phương án khả thi" có ưu nhược điểm; (e) "Đề xuất".')

H3('Bối cảnh chung của hai bài báo')

P('NCS Công Thị Hằng đang triển khai luận án tiến sĩ về rối loạn lo '
  'âu ở học sinh trung học cơ sở Hà Nội với bộ dữ liệu 1.352 học '
  'sinh, đo bằng tám thang chuẩn quốc tế. Để mở rộng giá trị khoa '
  'học của bộ dữ liệu này, nhóm dự kiến chiết xuất hai bài báo quốc '
  'tế:')

B('Bài Q1 (chất lượng Quý 1, tạp chí có chỉ số ảnh hưởng cao) gửi '
  'tới BMC Psychiatry — tập trung CƠ CHẾ tác động: phân tích bằng '
  'mô hình phương trình cấu trúc (Structural Equation Modeling, viết '
  'tắt SEM) — một phương pháp thống kê đa biến cho phép kiểm định '
  'đồng thời nhiều giả thuyết về quan hệ giữa các biến số.', 0)
B('Bài Q3 (chất lượng Quý 3) gửi tới PLOS ONE — tập trung MÔ TẢ chi '
  'tiết: tỷ lệ hiện mắc theo từng câu hỏi trong thang đo (item-level '
  'analysis), so sánh theo khối lớp và giới.', 0)

P('Sau khi em rà soát bản nháp Q1 phiên bản 3 (Draft_Q1_SongNgu_'
  'v3_01062026.docx) và đối chiếu với luận án chính, em xác định '
  'bốn câu hỏi mang tính chặn (BLOCKING) — tức là phải có câu trả '
  'lời thì mới hoàn thiện được bản thảo cuối cùng để gửi tạp chí. '
  'Bốn câu hỏi này không thể em tự quyết được vì liên quan đến: '
  '(a) dữ liệu chỉ NCS có; (b) chiến lược xuất bản của nhóm; '
  '(c) thủ tục đạo đức của trường.', italic=True)

H3('Bốn vấn đề trong file này ảnh hưởng tới')

B('Cấu trúc phương pháp được mô tả trong bản nháp', 0)
B('Cách trình bày kết quả phân tích', 0)
B('Khả năng vượt qua vòng sàng lọc đầu tiên của biên tập viên '
  '(editorial screening) — đây là vòng loại trực tiếp, không vào '
  'vòng phản biện nếu thiếu thông tin bắt buộc', 0)
B('Lộ trình nộp bài giữa hai tạp chí Q1 + Q3 — quyết định khi nào '
  'nộp bài nào trước', 0)


# ============================================================
H2('VẤN ĐỀ SỐ 1 — DỮ LIỆU PHỎNG VẤN ĐỊNH TÍNH (Q1-6)')

H3('Bối cảnh — Đang xảy ra điều gì?')
P('Bài Q1 dự kiến nộp tới tạp chí BMC Psychiatry (Q1 quốc tế, chỉ số '
  'ảnh hưởng JCR 2024 = 3,6) theo thiết kế nghiên cứu hỗn hợp song song '
  'hội tụ (convergent-parallel mixed-methods). Để dễ hiểu, thiết kế '
  'này nghĩa là nhóm nghiên cứu thực hiện SONG SONG hai phần tách biệt '
  'rồi kết hợp lại:')

B('Phần định lượng: phát phiếu khảo sát cho 1.352 học sinh, dùng các '
  'thang đo chuẩn để thu điểm số, sau đó chạy mô hình SEM để tìm các '
  'đường dẫn tác động.', 0)
B('Phần định tính: phỏng vấn trực tiếp một nhóm nhỏ học sinh khoảng '
  '30-45 phút mỗi em, để hiểu CÂU CHUYỆN sống động đằng sau các con '
  'số — vì sao một em điểm lo âu cao, áp lực gì cụ thể, vai trò gia '
  'đình thế nào.', 0)
B('Tích hợp: dùng ma trận joint-display để đặt cạnh nhau từng đường '
  'dẫn SEM (định lượng) và từng chủ đề định tính tương ứng. Ví dụ '
  'cụ thể: nếu SEM cho thấy "hỗ trợ cha mẹ" không có tác động bảo '
  'vệ đối với lo âu (β = -0,015, không có ý nghĩa), thì phần định '
  'tính phải tìm chủ đề giải thích vì sao — chẳng hạn "Im lặng '
  'trong gia đình" (cha mẹ không biết con đang lo âu vì con không '
  'dám chia sẻ). Nếu hai bên trùng khớp gọi là HỘI TỤ, lệch nhau '
  'gọi là PHÂN KỲ.', 0)

P('Trong bản nháp Q1 phiên bản 3 hiện tại, mục 2.3 (Phương pháp – '
  'phỏng vấn định tính) và mục 3.7 (Kết quả – tích hợp hỗn hợp) đang '
  'để chỗ trống có đánh dấu chờ NCS xác nhận thông tin chính thức.')

H3('Tại sao quan trọng?')

P('Reviewer của BMC Psychiatry không chấp nhận thiết kế hỗn hợp '
  '"trên giấy" — họ yêu cầu CHỨNG MINH thực hiện đầy đủ qua ba bằng '
  'chứng cụ thể: (a) sơ đồ chọn mẫu rõ ràng để biết NCS đã chọn '
  'những em nào, theo tiêu chí gì; (b) hệ số tin cậy giữa các mã '
  'hóa viên (Cohen κ) — để đảm bảo các chủ đề định tính không phải '
  'do một người tự nghĩ ra mà có nhiều người độc lập đồng thuận; '
  '(c) ma trận joint-display cho thấy cách tích hợp thực sự xảy ra.')

P('Nếu thiếu một trong ba bằng chứng này, kịch bản xấu nhất là bài '
  'bị reject ngay vòng đầu với lý do "qualitative methods inadequately '
  'reported" (báo cáo phương pháp định tính không đầy đủ). Đây là '
  'một trong những lý do phổ biến nhất khiến mixed-methods papers bị '
  'từ chối ở tạp chí Q1.', italic=True)

P('Trong bản nháp Q1 phiên bản 3 hiện tại, mục 2.3 (Phương pháp – '
  'phỏng vấn định tính) và mục 3.7 (Kết quả – tích hợp hỗn hợp) đang '
  'để chỗ trống có đánh dấu chờ NCS xác nhận thông tin chính thức.')

H3('Năm thông tin NCS cần xác nhận')
B('Số lượng người tham gia phỏng vấn (n = ?). Y văn quốc tế thường '
  'yêu cầu tối thiểu 12-15 người cho phân tích chủ đề chất lượng tốt, '
  'có thể đạt "saturation" (bão hòa) khoảng n = 20-25.', 0)
B('Chiến lược chọn mẫu có chủ đích — tức là không chọn ngẫu nhiên '
  'mà có tiêu chí cụ thể. Câu hỏi: NCS đã phân tầng theo mức lo âu '
  '(chọn cao – trung bình – thấp dựa trên điểm RCADS)? Theo giới? '
  'Theo khối lớp? Lựa chọn này phải được mô tả rõ.', 0)
B('Tình trạng tư liệu: ghi âm đã thực hiện chưa? Đã gỡ băng (chuyển '
  'audio sang văn bản) chưa? Tổng số trang transcript khoảng bao '
  'nhiêu? Trung bình một phỏng vấn 30-45 phút cho 15-25 trang văn bản.',
  0)
B('Hệ số đồng thuận liên mã hóa (intercoder reliability) Cohen κ — '
  'là chỉ số đo lường mức độ hai (hoặc ba) người mã hóa độc lập đồng '
  'ý với nhau khi gán cùng một chủ đề cho cùng một đoạn văn. Giá '
  'trị từ -1 đến 1. κ ≥ 0,60 là chấp nhận được, κ ≥ 0,80 là tốt. '
  'NCS đã tính chưa? Số mã hóa viên độc lập? Giá trị κ?', 0)
B('Khung mã hóa chủ đề: đã thiết lập theo quy trình sáu bước của '
  'Braun & Clarke (2006) [2] chưa? Đây là quy trình chuẩn quốc tế '
  'cho thematic analysis. Số chủ đề chính đã trích xuất? (Thường '
  '3-7 chủ đề là vừa phải cho một nghiên cứu hỗn hợp).', 0)

H3('Em đề xuất')
REC('NCS gửi cho em ba tài liệu: (a) sơ đồ chọn mẫu có chủ đích kèm '
    'bảng số người tham gia theo phân tầng; (b) ít nhất ba transcript '
    'mẫu để em rà soát chất lượng; (c) bảng kết quả Cohen κ kèm danh '
    'sách 3-5 chủ đề chính đã trích xuất. Em cần khoảng 7-10 ngày làm '
    'việc sau khi nhận để hoàn thiện mục 2.3 và 3.7.')

WARN('Nếu phần định tính chưa thực hiện đầy đủ, em đề xuất xem xét '
     'chuyển bài Q1 sang thiết kế thuần định lượng — chỉ giữ SEM tích '
     'hợp và bỏ phần hỗn hợp. Lựa chọn này giúp bài đi nhanh hơn nhưng '
     'mất đi điểm mới về cơ chế văn hóa.')


# ============================================================
H2('VẤN ĐỀ SỐ 2 — MÔ HÌNH SEM TÍCH HỢP (Q1-8)')

H3('Bối cảnh — Đang xảy ra điều gì?')

P('Để dễ hình dung, hãy tưởng tượng một sơ đồ với 7 ô bên trái (ba '
  'yếu tố nguy cơ: áp lực học tập, nghiện điện thoại, bắt nạt học '
  'đường + bốn yếu tố bảo vệ: gắn bó trường học, hỗ trợ cha mẹ, hỗ '
  'trợ bạn bè, tự trọng) và 3 ô bên phải (ba phân loại rối loạn lo '
  'âu: lo âu tổng quát GAD, lo âu chia ly SAD, lo âu xã hội SocAD). '
  'Mỗi mũi tên từ một ô trái sang một ô phải là một "đường dẫn" '
  '(path). Tổng cộng có 7 × 3 = 21 mũi tên.')

P('Tiêu đề và phần tóm tắt bài Q1 hiện tại đang khẳng định nghiên cứu '
  'kiểm định ĐỒNG THỜI cả 21 đường dẫn này trong một mô hình thống '
  'nhất — đây là cách diễn đạt "integrated risk-protective SEM" '
  '(SEM tích hợp). Lợi thế của mô hình tích hợp đồng thời là nó cho '
  'thấy ĐỘ ĐÓNG GÓP RIÊNG của mỗi yếu tố sau khi đã kiểm soát ảnh '
  'hưởng của các yếu tố khác — gần với thực tế hơn so với đo từng '
  'cặp một.')

P('Sau khi em rà soát luận án tiến sĩ chính của NCS (phiên bản '
  'FixCoping_28052026), thực tế các bảng số liệu thể hiện:')

B('Bảng 27, 30, 33: SEM riêng từng yếu tố nguy cơ → từng phân loại '
  'lo âu (mô hình hai biến – bivariate)', 0)
B('Bảng 36, 39, 42, 45: SEM riêng từng yếu tố bảo vệ → từng phân '
  'loại lo âu', 0)
B('Bảng 47: Mô hình tổng hợp với hai biến tiềm ẩn bậc cao — YTNC '
  '(tổng hợp ba yếu tố nguy cơ) và YTBV (tổng hợp bốn yếu tố bảo '
  'vệ) — dự báo tổng RLLA, với R² = 0,598', 0)

P('Tức là luận án KHÔNG chạy mô hình tích hợp 7 biến dự báo × 3 kết '
  'quả (21 đường dẫn) đồng thời như bản nháp Q1 hiện tại đang ngụ ý. '
  'Đây là một mâu thuẫn nội tại giữa nhan đề bài báo và phương pháp '
  'phân tích thực tế.', italic=True)

H3('Tại sao quan trọng?')

P('Reviewer của BMC Psychiatry hầu hết là chuyên gia thống kê SEM. '
  'Khi đọc tiêu đề "integrated SEM" họ sẽ kỳ vọng thấy một sơ đồ duy '
  'nhất với 21 mũi tên ước lượng đồng thời và một giá trị R² duy '
  'nhất cho mô hình tích hợp. Nếu phần Methods chỉ cung cấp các mô '
  'hình bivariate riêng lẻ, reviewer sẽ flag là "title-methods '
  'mismatch" (tiêu đề không khớp phương pháp). Đây là một trong các '
  'tình huống reject phổ biến nhất.')

P('Trong trường hợp này, em có ba phương án để giải quyết. Mỗi '
  'phương án có ưu nhược điểm khác nhau về tính trung thực với dữ '
  'liệu, thời gian thực hiện, và chất lượng bài cuối cùng.',
  italic=True)

H3('Ba phương án khả thi')

H3('Phương án A — Điều chỉnh cách diễn đạt (recommend)')
P('Giữ nguyên các phân tích đã có. Tái cấu trúc câu chuyện thành: '
  '"chúng tôi báo cáo các mô hình SEM theo từng phân loại (bivariate) '
  'cho mỗi cặp yếu tố – kết quả, sau đó tổng hợp thành mô hình bậc '
  'cao YTNC + YTBV → tổng RLLA (R² = 0,598)". Cách diễn đạt này trung '
  'thực với dữ liệu LA và vẫn truyền tải được khái niệm "tích hợp". '
  'Không cần phân tích lại.')

P('Ưu điểm: Sẵn sàng cho bản nháp v4 trong vòng 1-2 ngày. Không cần '
  'truy cập lại bộ dữ liệu thô.', indent=False)
P('Nhược điểm: Phần "tích hợp" yếu hơn so với mô hình SEM 21 đường '
  'dẫn đầy đủ. Reviewer có thể yêu cầu giải thích.', indent=False)

H3('Phương án B — Tái phân tích SEM tích hợp đầy đủ')
P('Chạy lại AMOS với mô hình đơn lẻ: 7 biến quan sát (hoặc 7 biến '
  'tiềm ẩn nếu giữ cấu trúc đo lường) → 3 biến kết quả tiềm ẩn (GAD, '
  'SAD, SocAD) → tổng cộng 21 đường dẫn ước lượng đồng thời. Báo cáo '
  '1 giá trị R² đại diện cho cấu trúc tích hợp.')

P('Ưu điểm: Đúng nghĩa "integrated SEM" theo y văn quốc tế. Trùng '
  'khớp với tiêu đề. Có thể tiết lộ thêm mẫu hình tương tác giữa các '
  'yếu tố mà mô hình bivariate không phát hiện được.', indent=False)
P('Nhược điểm: Cần truy cập bộ dữ liệu thô (SPSS/AMOS file). Em ước '
  'tính 1-2 ngày phân tích + 2-3 ngày báo cáo lại. Có rủi ro mô hình '
  'không hội tụ nếu cỡ mẫu chia trên 21 đường dẫn không đủ — quy tắc '
  'tối thiểu là 10 quan sát / đường dẫn, tức cần ≥ 210 trường hợp '
  '(mẫu 1.352 thừa).', indent=False)

H3('Phương án C — Thu hẹp phạm vi bài')
P('Giảm bài Q1 xuống còn 3 yếu tố nguy cơ × 3 phân loại lo âu (9 '
  'đường dẫn) HOẶC 4 yếu tố bảo vệ × 3 phân loại (12 đường dẫn). '
  'Tách yếu tố còn lại sang bài thứ ba.')

P('Ưu điểm: Đơn giản, nhanh, an toàn. Mỗi bài gọn về mặt phạm vi.',
  indent=False)
P('Nhược điểm: Mất đi điểm mới về "tích hợp đồng thời". Phải đầu tư '
  'thêm thời gian viết bài thứ ba.', indent=False)

REC('Phương án A cho bản nháp v4. Phương án B chỉ tiến hành nếu '
    'reviewer của BMC Psychiatry yêu cầu cụ thể trong vòng phản biện '
    'đầu tiên. Phương án C chỉ áp dụng nếu thầy có chiến lược nộp '
    'thành 3 bài độc lập.')


# ============================================================
H2('VẤN ĐỀ SỐ 3 — LETTER CHẤP THUẬN ĐẠO ĐỨC HNUE (Q3-6)')

H3('Bối cảnh — Đang xảy ra điều gì?')

P('Khi nghiên cứu khoa học có sự tham gia của con người (đặc biệt là '
  'trẻ em và thanh thiếu niên), luật quốc tế yêu cầu phải có sự phê '
  'duyệt trước của một Hội đồng Đạo đức (Institutional Review Board, '
  'viết tắt IRB hoặc Ethics Committee). Mục đích là đảm bảo: (a) học '
  'sinh tham gia tự nguyện và có sự đồng ý của cha mẹ; (b) thông tin '
  'cá nhân được bảo vệ; (c) nghiên cứu không gây hại tâm lý; (d) học '
  'sinh được quyền rút khỏi nghiên cứu bất cứ lúc nào.')

P('Cả hai tạp chí mục tiêu (BMC Psychiatry và PLOS ONE) đều BẮT BUỘC '
  'tác giả khai báo SỐ QUYẾT ĐỊNH và NGÀY BAN HÀNH cụ thể của Hội đồng '
  'Đạo đức trong phần Phương pháp [3] [4]. Họ không chấp nhận lời '
  'tuyên bố chung chung "được phê duyệt bởi cơ quan có thẩm quyền". '
  'Phải có số văn bản như "Quyết định số 123/QĐ-ĐHSPHN ngày 15/03/2025" '
  'mới hợp lệ.')

P('Đây là điều kiện sàng lọc đầu tiên — bài không có thông tin IRB '
  'sẽ bị biên tập viên trả về ngay, không vào vòng phản biện. Thời '
  'gian xử lý chỉ 1-2 tuần.', italic=True)

P('Mục 2.5 Đạo đức trong bản nháp Q1 hiện tại đang để chỗ trống chờ '
  'thông tin chính thức.')

H3('Hai trường hợp cần NCS xác nhận')

H3('Trường hợp 1 — Đã có letter chính thức')
B('Số quyết định: ___', 0)
B('Ngày ban hành (dd/mm/yyyy): ___', 0)
B('Tên đầy đủ của hội đồng (ví dụ: Hội đồng Đạo đức Trường Đại học '
  'Sư phạm Hà Nội): ___', 0)
B('Hiệu lực: vĩnh viễn / có thời hạn (ghi rõ ngày hết hạn): ___', 0)
B('NCS gửi em bản scan/photo letter để em lưu kèm bản nháp', 0)

H3('Trường hợp 2 — Chưa có letter chính thức')
P('NCS cần liên hệ Phòng Quản lý Khoa học và Hợp tác Quốc tế của '
  'Trường Đại học Sư phạm Hà Nội để xin phê duyệt đạo đức truy hồi '
  '(retroactive IRB approval). Thời gian xử lý thường 2-4 tuần. Một '
  'số trường có Hội đồng Đạo đức cấp khoa (Khoa Tâm lý-Giáo dục) — '
  'NCS có thể xin tại đây nếu nhanh hơn.')

H3('Rủi ro')
WARN('PLOS ONE và BMC Psychiatry KHÔNG chấp nhận lời khẳng định "nghiên '
     'cứu được phê duyệt bởi cơ quan có thẩm quyền" mà không có số '
     'quyết định cụ thể. Bài không có IRB number sẽ bị reject ngay '
     'trong vòng 1-2 tuần ở vòng sàng lọc biên tập.')

REC('NCS xác nhận trạng thái letter trong tuần này. Nếu chưa có, '
    'song song hai việc: (a) xin letter chính thức qua Phòng QLKH '
    'HNUE; (b) tiếp tục hoàn thiện bản nháp Q1 + Q3 trong khi chờ. '
    'Bản nháp chỉ thực sự sẵn sàng nộp khi có IRB number.')


# ============================================================
H2('VẤN ĐỀ SỐ 4 — CHIẾN LƯỢC NỘP BÀI Q1 ↔ Q3 (Q3-9)')

H3('Bối cảnh — Đang xảy ra điều gì?')

P('Khi cùng một bộ dữ liệu được dùng để viết NHIỀU bài báo, các tạp '
  'chí quốc tế đặt ra quy tắc nghiêm ngặt để tránh "salami slicing" '
  '— tức là chia một nghiên cứu thành nhiều bài nhỏ chỉ để tăng số '
  'lượng xuất bản. Quy định của Ủy ban Quốc tế các Biên tập viên '
  'Tạp chí Y học (ICMJE) [5] yêu cầu: (a) mỗi bài phải trả lời một '
  'câu hỏi nghiên cứu RIÊNG BIỆT; (b) phải khai báo rõ với biên tập '
  'viên rằng có bài liên quan đang ở giai đoạn nào (nộp, đang phản '
  'biện, đã chấp nhận); (c) tránh trùng lặp câu chữ giữa các bài.')

P('Cả hai bài Q1 (BMC Psychiatry, tập trung phân tích cơ chế SEM + '
  'phương pháp hỗn hợp) và Q3 (PLOS ONE, tập trung mô tả chi tiết '
  'mức độ mục — item-level) sử dụng CÙNG bộ dữ liệu 1.352 học sinh '
  'trung học cơ sở Hà Nội. Đây là nguồn rủi ro cho hai vấn đề: '
  '(a) trùng lặp xuất bản (redundant publication) khi cả hai bài '
  'báo cáo cùng những con số quan trọng; (b) tự đạo văn '
  '(self-plagiarism) khi cùng đoạn văn xuất hiện trong cả hai bản '
  'nháp.')

H3('Tại sao quan trọng?')

P('Phần mềm phát hiện đạo văn (như Crossref Similarity Check, '
  'Turnitin) so sánh bản nháp mới với CSDL các bài đã xuất bản — '
  'bao gồm cả bài của chính tác giả. Nếu bài Q3 trùng > 20% câu chữ '
  'với bài Q1 đã xuất bản, biên tập viên PLOS ONE sẽ flag tự đạo '
  'văn và yêu cầu sửa hoặc reject. Vì vậy chiến lược nộp bài hai '
  'tạp chí cần được lên kế hoạch CHỦ ĐỘNG ngay từ đầu.', italic=True)

H3('Ba phương án nộp bài')
P('Mỗi phương án có sự đánh đổi giữa TỐC ĐỘ và AN TOÀN. Phương án A '
  'an toàn nhất nhưng chậm. Phương án B nhanh nhất nhưng rủi ro. '
  'Phương án C là biện pháp chống tự đạo văn bắt buộc áp dụng dù '
  'chọn A hay B.', italic=True)

H3('Phương án A — Tuần tự (sequential)')
P('Nộp Q1 trước → chờ vòng phản biện và quyết định chấp nhận (3-6 '
  'tháng) → nộp Q3 sau, trong đó trích dẫn rõ Q1 đã chấp nhận: "for '
  'the integrated SEM analysis see [Cong & Dao 2026 BMC Psychiatry]".')

P('Ưu điểm: An toàn nhất. Mỗi bài có vị trí riêng. Khả năng được '
  'biên tập viên cùng nhóm tạp chí flag overlap thấp.', indent=False)
P('Nhược điểm: Tổng thời gian từ nộp Q1 đến chấp nhận Q3 khoảng '
  '8-12 tháng. Mất cơ hội nếu NCS cần xuất bản nhanh để hoàn thành '
  'yêu cầu bảo vệ luận án.', indent=False)

H3('Phương án B — Song song (companion papers)')
P('Nộp Q1 và Q3 cùng lúc. Thêm chú thích chân trang trong mục Phương '
  'pháp của cả hai bài: "This study is one of two companion papers '
  'from the same dataset; the [other paper] is under review at '
  '[journal]". Một số tạp chí (đặc biệt PLOS family) chấp nhận chuẩn '
  'này nếu hai bài bổ sung lẫn nhau về mặt câu hỏi nghiên cứu.')

P('Ưu điểm: Tổng thời gian 4-6 tháng. Nhanh hơn đáng kể.', indent=False)
P('Nhược điểm: Rủi ro biên tập viên cả hai tạp chí phát hiện overlap '
  'và yêu cầu hợp nhất hoặc rút bài. BMC Psychiatry và PLOS ONE thuộc '
  'hai publisher khác nhau (Springer Nature và PLOS) nhưng có thể có '
  'shared reviewers trong cùng lĩnh vực adolescent mental health Việt '
  'Nam.', indent=False)

H3('Phương án C — Tách trọng tâm hoàn toàn (luôn cần dù chọn A hay B)')
P('Đây là biện pháp bắt buộc, không loại trừ với A hoặc B:')
B('Q1 tập trung CƠ CHẾ: SEM tích hợp, multi-group invariance, tích '
  'hợp định tính. KHÔNG báo cáo tỷ lệ mục theo từng câu RCADS.', 0)
B('Q3 tập trung MÔ TẢ: tỷ lệ hiện mắc theo từng mục của RCADS '
  '(item-level), so sánh theo khối lớp, theo giới. KHÔNG báo cáo β '
  'SEM của Q1.', 0)

REC('Phương án A + C nếu thầy ưu tiên an toàn (recommend). Phương án '
    'B + C nếu NCS cần xuất bản nhanh để phục vụ bảo vệ luận án. Em '
    'sẵn sàng viết bản nháp Q3 song song với Q1 nếu thầy chọn B.')


# ============================================================
H2('LỊCH TRÌNH ĐỀ XUẤT')

P('Em đề xuất lịch trình các bước tiếp theo (giả định phương án A + '
  'C được chọn):')

t = d.add_table(rows=1, cols=3); t.style = 'Light Grid Accent 1'; t.autofit = False
hdr = t.rows[0].cells
for i, h in enumerate(['Tuần', 'Việc cần làm', 'Người chịu trách nhiệm']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)
rows = [
    ('Tuần 1', 'NCS xác nhận 4 vấn đề BLOCKING. Gửi em thông tin '
     'định tính + scan letter IRB (nếu có)', 'NCS + thầy HD'),
    ('Tuần 2', 'Em tích hợp dữ liệu định tính + IRB number vào bản '
     'nháp Q1 v4. Sửa Section 2.3, 2.5, 3.7', 'Em (Hải)'),
    ('Tuần 3', 'Đồng tác giả Nguyễn Minh Đức + thầy HD đọc + góp ý '
     'v4. Em chỉnh sửa theo phản hồi', 'Tất cả'),
    ('Tuần 4', 'Bản nháp Q1 v5 sẵn sàng nộp BMC Psychiatry. Em soạn '
     'cover letter + suggested reviewers', 'Em + NCS'),
    ('Tuần 4', 'Em bắt đầu soạn bản nháp Q3 (item-level descriptive)',
     'Em (Hải)'),
    ('Tuần 5-6', 'Hoàn thiện Q3. Submit Q1 tới BMC Psychiatry',
     'NCS submit'),
    ('Tuần 7+', 'Chờ phản biện Q1 (3-6 tháng). Trong thời gian này: '
     'hoàn thiện Q3, chờ quyết định Q1', 'Tất cả'),
]
for row_data in rows:
    row = t.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
set_col_widths(t, [2.5, 9.5, 4.0])


# ============================================================
H2('TÓM TẮT 4 CÂU HỎI CHÍNH CHO THẦY + NCS')

P('Để thuận tiện cho thầy hồi đáp, em tóm tắt bốn câu hỏi cốt lõi:')

q_table_data = [
    ['Mã', 'Câu hỏi', 'Người trả lời'],
    ['Q1-6', 'Dữ liệu phỏng vấn định tính: n, sampling, transcripts, '
     'Cohen κ, khung mã hóa đã có ở mức nào?', 'NCS'],
    ['Q1-8', 'Phương án A (điều chỉnh diễn đạt), B (tái phân tích '
     'SEM 21 đường dẫn), hay C (thu hẹp bài) cho mô hình tích hợp?',
     'Thầy HD + NCS'],
    ['Q3-6', 'Letter chấp thuận đạo đức HNUE: đã có chính thức chưa? '
     'Số quyết định + ngày?', 'NCS'],
    ['Q3-9', 'Chiến lược nộp bài: A (tuần tự, 8-12 tháng), B '
     '(song song, 4-6 tháng), hay C (tách trọng tâm)?',
     'Thầy HD + NCS'],
]

t2 = d.add_table(rows=1, cols=3); t2.style = 'Light Grid Accent 1'; t2.autofit = False
hdr = t2.rows[0].cells
for i, h in enumerate(q_table_data[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)
for row_data in q_table_data[1:]:
    row = t2.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
set_col_widths(t2, [1.5, 11.0, 3.5])


# ============================================================
H2('THAM KHẢO')

refs = [
    '[1] Creswell, J. W., & Plano Clark, V. L. (2018). Designing and '
    'Conducting Mixed Methods Research (3rd ed.). Thousand Oaks, CA: '
    'Sage Publications. ISBN 978-1-4833-4437-9.',

    '[2] Braun, V., & Clarke, V. (2006). Using thematic analysis in '
    'psychology. Qualitative Research in Psychology, 3(2), 77–101. '
    'DOI: 10.1191/1478088706qp063oa.',

    '[3] BMC Psychiatry. Editorial Policies: Ethics approval and '
    'consent to participate. Truy cập: https://bmcpsychiatry.bio'
    'medcentral.com/submission-guidelines/preparing-your-manuscript',

    '[4] PLOS ONE. Submission Guidelines: Ethics Statement '
    '(Human Subjects Research). Truy cập: https://journals.plos.org/'
    'plosone/s/submission-guidelines#loc-ethics-statement',

    '[5] International Committee of Medical Journal Editors (ICMJE). '
    'Recommendations for the Conduct, Reporting, Editing, and '
    'Publication of Scholarly Work in Medical Journals. Section II.D: '
    'Overlapping Publications. Truy cập: https://www.icmje.org/'
    'recommendations/',

    '[6] Luận án TS NCS Công Thị Hằng, phiên bản 01_LuanAn_v3_1_'
    'FixCoping_28052026.docx. Các bảng đã đối chiếu: Bảng 27 (ALHT '
    '→ phân loại lo âu), Bảng 30 (NĐT), Bảng 33 (BNHĐ), Bảng 36 '
    '(GBTH), Bảng 39 (TTr), Bảng 42 (HTCM), Bảng 45 (HTBB), Bảng 47 '
    '(YTNC + YTBV → RLLA).',

    '[7] Outline_Q1_v3_01062026.docx — bản đề cương Q1 phiên bản 3 '
    'có 8 vấn đề HIGH/MEDIUM đã sửa kèm 2 BLOCKING đánh dấu placeholder. '
    'Tham chiếu nội bộ.',

    '[8] Memory `project_BLOCKING_Q1Q3_questions.md` (em ghi nhận '
    '01/06/2026 trong hệ thống nhớ phiên hội thoại). Tham chiếu nội bộ.',

    '[9] Draft_Q1_SongNgu_v3_01062026.docx — bản nháp song ngữ '
    'Anh – Việt v3 với 7 chỗ chờ [TBD] đánh dấu rõ ràng. Tham chiếu '
    'nội bộ.',
]

for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# ============================================================
H2('LIÊN HỆ')

P('Em sẵn sàng làm việc trực tiếp với thầy và NCS để giải quyết bốn '
  'vấn đề trên. Mọi thắc mắc xin gửi qua kênh trao đổi hiện tại. Sau '
  'khi nhận được phản hồi cho cả bốn câu hỏi, em sẽ tích hợp vào bản '
  'nháp Q1 phiên bản 4 trong vòng 7-10 ngày làm việc.', indent=False)

P('Trân trọng cảm ơn thầy và NCS.', italic=True, indent=False)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
