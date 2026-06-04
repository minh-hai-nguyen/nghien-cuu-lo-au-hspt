"""Build doc tra loi: tac gia khao sat coping 2007 Mid-Atlantic = Herres & Ohannessian 2015"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Tac_gia_khao_sat_coping_2007_Herres_Ohannessian_2015.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(12); r.italic = italic
    return p

# =====================================================
# TIÊU ĐỀ
# =====================================================
H('Tác giả khảo sát coping 2007 Mid-Atlantic Hoa Kỳ', level=1)
H('— bài "Các hồ sơ ứng phó của vị thành niên phân biệt báo cáo triệu chứng trầm cảm và lo âu" —', level=2)

# =====================================================
# CÂU HỎI (xanh)
# =====================================================
H('Câu hỏi của thầy', level=2)
para_blue(
    'Cái này trong bài "Ứng phó", ai là tác giả của khảo sát 2007 này? '
    '"Khảo sát cắt ngang (cross-sectional survey) tại 7 trường trung học phổ thông '
    'công lập vùng Mid-Atlantic Hoa Kỳ vào mùa xuân 2007." '
    'Bài: "CÁC HỒ SƠ ỨNG PHÓ CỦA VỊ THÀNH NIÊN PHÂN BIỆT BÁO CÁO TRIỆU CHỨNG TRẦM CẢM VÀ LO ÂU"'
)

# =====================================================
# CÂU TRẢ LỜI (xanh, gom 1 chỗ — vì đây là câu hỏi đơn giản)
# =====================================================
H('1. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Tác giả:', bold=True)
para_blue(
    'Joanna Herres + Christine McCauley Ohannessian (2015).'
)

para_blue('Tên bài đầy đủ:', bold=True)
para_blue(
    '"Adolescent coping profiles differentiate reports of depression and anxiety '
    'symptoms" — "Các hồ sơ ứng phó của vị thành niên phân biệt các báo cáo về '
    'triệu chứng trầm cảm và lo âu".'
)

para_blue('Xuất bản:', bold=True)
bullet_blue('Tạp chí: Journal of Affective Disorders (Elsevier).')
bullet_blue('Tập 186, trang 312–319, số đặc biệt tháng 8/2015.')
bullet_blue('DOI: 10.1016/j.jad.2015.07.031')
bullet_blue('PubMed Central: PMC4565746')

para_blue('Đặc điểm khảo sát + mẫu:', bold=True)
bullet_blue('Thiết kế: Khảo sát cắt ngang (cross-sectional survey).')
bullet_blue('Địa điểm: 7 trường trung học phổ thông công lập vùng Mid-Atlantic Hoa Kỳ.')
bullet_blue('Thời gian: mùa xuân 2007 (spring 2007).')
bullet_blue('Cỡ mẫu: N = 982 học sinh THPT (54% nữ; 46% nam).')
bullet_blue('Tuổi trung bình: 16,09 (SD = 0,68); lớp 10–11.')

para_blue('Lưu ý quan trọng:', bold=True)
para_blue(
    'Khảo sát thực hiện năm 2007 nhưng bài báo CÔNG BỐ năm 2015 — khoảng cách 8 năm. '
    'Đây là điều bình thường: nhiều dataset lớn được phân tích nhiều lần qua nhiều năm. '
    'Khi trích dẫn vào báo cáo CTH, ghi năm công bố (2015), không phải năm thu thập '
    'dữ liệu (2007).'
)

# =====================================================
# BACKGROUND BỔ SUNG (đen)
# =====================================================
H('2. Thông tin bổ sung về tác giả', level=2)

para_black('Christine McCauley Ohannessian:', bold=True)
bullet_black(
    'PhD. Theo y văn (PubMed search), là PI nổi tiếng về tâm lý học vị thành niên, đặc '
    'biệt về ứng phó, gia đình, sử dụng chất gây nghiện; có nhiều công bố về adolescent '
    'coping/family functioning.', italic=True
)
bullet_black(
    '⚠ Affiliation cụ thể (UConn Health hay nơi khác) — em CHƯA verify từ web. Thầy có '
    'thể tra qua https://pubmed.ncbi.nlm.nih.gov bằng query "Ohannessian CM[Author]" '
    'để xác nhận affiliation hiện tại.', italic=True
)

para_black('Joanna Herres:', bold=True)
bullet_black(
    'PhD. First author của bài này (2015) — thường là PhD student hoặc postdoc của '
    'PI Ohannessian khi đó. Hiện vai trò trong bài hợp tác Herres + Ohannessian sau 2015 '
    'cho thấy duy trì hợp tác nghiên cứu.', italic=True
)
bullet_black(
    '⚠ Affiliation hiện tại — em CHƯA verify từ web. Thầy tra PubMed bằng query '
    '"Herres J[Author] coping" để xác nhận.', italic=True
)

H('3. Vì sao bài này quan trọng cho đề tài lo âu HS THCS/THPT của thầy?', level=2)
bullet_black('Phân biệt rõ "coping profiles" (mẫu hình ứng phó) giữa nhóm có TRẦM CẢM vs nhóm có LO ÂU vs nhóm có CẢ HAI — quan trọng cho can thiệp phân biệt.')
bullet_black('Mẫu lớn (N=982) + tuổi 16 (lớp 10-11) — phù hợp với đối tượng THPT VN của thầy.')
bullet_black('Thang đo: Children\'s Depression Inventory (CDI) + State-Trait Anxiety Inventory for Children (STAIC) + COPE Inventory — bộ 3 thang chuẩn.')
bullet_black('Khu vực Mid-Atlantic = các bang Pennsylvania, New Jersey, Delaware, Maryland — phổ điểm SES đa dạng, có thể so sánh với HS THPT VN.')

H('4. Cách trích dẫn chuẩn APA 7', level=2)
para_black(
    'Herres, J., & Ohannessian, C. M. (2015). Adolescent coping profiles differentiate '
    'reports of depression and anxiety symptoms. Journal of Affective Disorders, 186, '
    '312–319. https://doi.org/10.1016/j.jad.2015.07.031', italic=True
)

H('5. Phụ lục — Tham chiếu trong corpus dự án', level=2)
bullet_black(
    'Bài Herres & Ohannessian (2015) đã được liệt kê trong file '
    '01_Bao-cao/01_Gửi H_DANH MỤC TÀI LIỆU THAM KHẢO.docx (đoạn 85).', italic=True
)
bullet_black(
    'Bài này CHƯA được canonical hóa vào DB chính (chưa có mã VNxxx/QTxxx). '
    'Nếu thầy muốn em dịch + tóm tắt + cấp mã canonical, em có thể làm.', italic=True
)
bullet_black(
    'Liên hệ với corpus: bài QT066 Murphy 2024 (peer support) + QT067 Pascoe 2020 '
    '(academic stress) đều có thảo luận về coping styles ở vị thành niên — có thể '
    'cross-reference khi thầy viết.', italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
