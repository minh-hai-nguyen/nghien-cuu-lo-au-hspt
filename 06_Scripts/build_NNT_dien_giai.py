"""Build doc: NNT dien giai cu the nhu the nao + vi sao cang nho cang hieu qua."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

# TITLE
title = d.add_heading('NNT (Number Needed to Treat): Diễn giải và vì sao càng nhỏ càng hiệu quả?', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Hướng dẫn đọc một trong những chỉ số quan trọng nhất trong evidence-based medicine')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# META
meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
mp = cell.paragraphs[0]
mp.add_run('Câu hỏi gốc của thầy: ').bold = True
mp.add_run('"NNT diễn giải cụ thể như thế nào? Vì sao chỉ số này càng nhỏ thì biện pháp can thiệp càng hiệu quả?"')
d.add_paragraph()

# =========================================================
# PHẦN 1 — ĐỊNH NGHĨA + DIỄN GIẢI
# =========================================================
add_heading('1. NNT là gì?', level=1)

add_para('NNT (Number Needed to Treat) = Số người cần điều trị để có THÊM 1 người cải thiện so với không điều trị (hoặc so với nhóm đối chứng).', bold=True, color=(192, 0, 0))
d.add_paragraph()

add_para('Diễn giải bằng câu một dòng:', bold=True, size=12)
add_para('"Cứ điều trị NNT người, sẽ có THÊM 1 người đáp ứng nhờ can thiệp."')
d.add_paragraph()

# Example box — CAMS
cams_box = d.add_table(rows=1, cols=1)
cams_box.style = 'Table Grid'
cc = cams_box.rows[0].cells[0]
shade_cell(cc, 'DEEBF7')
cp = cc.paragraphs[0]
cr = cp.add_run('Ví dụ thực tế — CAMS (Walkup 2008) cho kết hợp CBT + Sertraline:\n')
cr.bold = True
cr.font.color.rgb = RGBColor(31, 73, 125)
cp.add_run('NNT = 3 → Cứ 3 trẻ em có rối loạn lo âu được điều trị bằng CBT + Sertraline, sẽ có THÊM 1 trẻ đáp ứng lâm sàng (giảm triệu chứng rõ rệt) so với nếu dùng giả dược. Hai trẻ còn lại có thể đã tự hồi phục hoặc vẫn không đáp ứng — nhưng 1 trẻ được "cứu" nhờ điều trị.')

d.add_paragraph()

# =========================================================
# PHẦN 2 — CÔNG THỨC
# =========================================================
add_heading('2. Công thức NNT', level=1)

add_para('Công thức:', bold=True)
add_para('NNT = 1 / ARR', bold=True, size=14, color=(192, 0, 0))
add_para('Trong đó:')
for it in [
    '• ARR = Absolute Risk Reduction (Giảm rủi ro tuyệt đối)',
    '• ARR = Tỷ lệ đáp ứng nhóm can thiệp − Tỷ lệ đáp ứng nhóm đối chứng',
    '• ARR tính bằng tỷ lệ (số thập phân 0–1), KHÔNG phải phần trăm',
]:
    p = d.add_paragraph(it)
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# Tính NNT cho CAMS
add_para('Tính NNT cho CAMS (combination CBT + Sertraline):', bold=True, size=12)

calc_tbl = d.add_table(rows=4, cols=2)
calc_tbl.style = 'Table Grid'
calcs = [
    ('Tỷ lệ đáp ứng nhóm CBT+Sertraline', '80,7 % = 0,807'),
    ('Tỷ lệ đáp ứng nhóm placebo', '23,7 % = 0,237'),
    ('ARR = 0,807 − 0,237', '0,57 (tức 57 %)'),
    ('NNT = 1 / 0,57', '≈ 1,75 → làm tròn lên thành 2 (hoặc 3 như Zugman 2024 ghi)'),
]
for i, (k, v) in enumerate(calcs):
    row = calc_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v

d.add_paragraph()

note_tbl = d.add_table(rows=1, cols=1)
note_tbl.style = 'Table Grid'
nc = note_tbl.rows[0].cells[0]
shade_cell(nc, 'FFF2CC')
np_ = nc.paragraphs[0]
nr = np_.add_run('Quy tắc làm tròn: ')
nr.bold = True
nr.font.color.rgb = RGBColor(191, 97, 14)
np_.add_run('NNT LUÔN làm tròn LÊN (không làm tròn xuống). Vì nếu NNT = 1,75 thì trên thực tế cần ít nhất 2 bệnh nhân để đảm bảo có 1 người đáp ứng thêm. Không thể "điều trị 1,75 người".')
d.add_paragraph()

# =========================================================
# PHẦN 3 — VÌ SAO CÀNG NHỎ CÀNG TỐT
# =========================================================
add_heading('3. Vì sao NNT càng NHỎ thì can thiệp càng HIỆU QUẢ?', level=1, color=(192, 80, 77))

add_para('Logic toán học đơn giản:', bold=True, size=12)
d.add_paragraph()

logic_tbl = d.add_table(rows=6, cols=2)
logic_tbl.style = 'Table Grid'
hdr = logic_tbl.rows[0]
for i, h in enumerate(['Tình huống', 'Ý nghĩa']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

logic_data = [
    ('NNT = 1 (lý tưởng, hiếm)',
     'Điều trị 1 người → chắc chắn có 1 người được cải thiện. Can thiệp cực kỳ hiệu quả, gần như "bắt buộc có tác dụng". Ít thấy trong y học — ví dụ insulin cho bệnh nhân đái tháo đường type 1.'),
    ('NNT = 3 (rất hiệu quả)',
     'Điều trị 3 người → 1 người cải thiện thêm. CAMS combination CBT+Sertraline nằm trong mức này. Là "gold standard" của psychotherapy + pharmacotherapy.'),
    ('NNT = 10 (hiệu quả vừa)',
     'Điều trị 10 người → 1 người cải thiện thêm. Chấp nhận được trong lâm sàng nếu can thiệp an toàn + rẻ. Nhiều SSRIs cho trầm cảm nhẹ nằm ở mức này.'),
    ('NNT = 50 (hiệu quả thấp)',
     'Điều trị 50 người → chỉ 1 người cải thiện thêm. Thường áp dụng cho can thiệp DỰ PHÒNG quần thể lớn (ví dụ aspirin phòng tim mạch). Nhân rộng cần tính cost-benefit kỹ.'),
    ('NNT = ∞ (vô hạn)',
     'Không có chênh lệch giữa nhóm can thiệp và đối chứng → ARR = 0 → NNT vô cùng. Can thiệp không có tác dụng.'),
]
for i, (k, v) in enumerate(logic_data):
    row = logic_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = v

d.add_paragraph()

# Analogy box
ana_tbl = d.add_table(rows=1, cols=1)
ana_tbl.style = 'Table Grid'
ac = ana_tbl.rows[0].cells[0]
shade_cell(ac, 'FFF8DC')
ap = ac.paragraphs[0]
ar = ap.add_run('Phép ẩn dụ dễ nhớ — "LƯỚI cá":\n')
ar.bold = True
ar.font.color.rgb = RGBColor(191, 97, 14)
ap.add_run('Hãy tưởng tượng NNT là số "con mồi" phải thả xuống lưới để BẮT được 1 con cá THÊM. Lưới hiệu quả (can thiệp tốt) → chỉ cần thả 3 con mồi là bắt được 1 con cá. Lưới kém → thả 50 con mồi mới bắt được 1 con. Vậy NNT càng NHỎ = lưới càng HIỆU QUẢ.')

d.add_paragraph()

# Intuition
add_para('Hiểu theo trực giác công thức:', bold=True, size=12)
intuit_tbl = d.add_table(rows=3, cols=2)
intuit_tbl.style = 'Table Grid'
intuits = [
    ('Công thức NNT = 1 / ARR',
     'ARR là hiệu suất tạo thêm người khỏi. ARR càng LỚN → hiệu suất càng cao → phép chia 1/ARR sẽ ra số càng NHỎ.'),
    ('ARR = 57% (như CAMS)',
     '1 / 0,57 ≈ 1,75 → NNT nhỏ → can thiệp mạnh. Một can thiệp tạo thêm 57/100 người khỏi là rất lớn trong medicine.'),
    ('ARR = 2% (nhiều can thiệp quần thể)',
     '1 / 0,02 = 50 → NNT = 50 → can thiệp yếu. Tạo thêm 2/100 người khỏi nghĩa là chỉ "vớt" được 2 người trên 100.'),
]
for i, (k, v) in enumerate(intuits):
    row = intuit_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v

d.add_paragraph()

# =========================================================
# PHẦN 4 — NGƯỠNG ĐỌC LÂM SÀNG
# =========================================================
add_heading('4. Ngưỡng đọc NNT trong lâm sàng (benchmark)', level=1)

bench_tbl = d.add_table(rows=6, cols=3)
bench_tbl.style = 'Table Grid'
hdr = bench_tbl.rows[0]
for i, h in enumerate(['NNT', 'Đánh giá', 'Ví dụ thực tế']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

bench = [
    ('1–3', 'XUẤT SẮC', 'Insulin cho đái tháo đường type 1 | CAMS combination CBT+SSRI cho lo âu trẻ em (NNT=3)'),
    ('4–10', 'RẤT TỐT', 'CBT đơn trị cho lo âu trẻ em (NNT≈4–6) | SSRIs cho trầm cảm nặng'),
    ('11–20', 'TỐT', 'SSRIs cho trầm cảm vừa | Thuốc huyết áp cho BN tăng huyết áp nặng'),
    ('21–50', 'TRUNG BÌNH', 'Statins dự phòng tim mạch thứ phát (NNT≈30–50) | Nhiều can thiệp dự phòng'),
    ('> 50', 'THẤP', 'Aspirin dự phòng tiên phát tim mạch (NNT≈100) | Can thiệp quần thể diện rộng'),
]
for i, (r_, k, ex) in enumerate(bench):
    row = bench_tbl.rows[i+1]
    row.cells[0].text = r_
    row.cells[1].text = k
    row.cells[2].text = ex
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
    for p in row.cells[1].paragraphs:
        for r in p.runs:
            r.bold = True

d.add_paragraph()

caveat = d.add_table(rows=1, cols=1)
caveat.style = 'Table Grid'
cv = caveat.rows[0].cells[0]
shade_cell(cv, 'FCE4D6')
cvp = cv.paragraphs[0]
cvr = cvp.add_run('⚠ Lưu ý: ')
cvr.bold = True
cvr.font.color.rgb = RGBColor(192, 0, 0)
cvp.add_run('NNT chỉ có ý nghĩa khi đi kèm: (1) thời gian theo dõi (ví dụ NNT=3 trong 12 tuần vs NNT=3 trong 1 năm rất khác nhau); (2) định nghĩa "response" (giảm 50% triệu chứng, hay remission hoàn toàn); (3) nhóm đối chứng là gì (placebo, waitlist, TAU). NNT không đứng độc lập — phải đọc cùng hoàn cảnh cụ thể.')

d.add_paragraph()

# =========================================================
# PHẦN 5 — VÍ DỤ TỪ DỰ ÁN
# =========================================================
add_heading('5. Ví dụ NNT từ các bài trong thư viện dự án', level=1)

proj_tbl = d.add_table(rows=5, cols=3)
proj_tbl.style = 'Table Grid'
hdr = proj_tbl.rows[0]
for i, h in enumerate(['Bài / can thiệp', 'NNT / response rate', 'Đọc như thế nào']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

projs = [
    ('Walkup 2008 CAMS — CBT+Sertraline vs Placebo',
     '80,7 % vs 23,7 % → NNT ≈ 2–3',
     'XUẤT SẮC. 2-3 trẻ điều trị để có 1 trẻ cải thiện thêm. Combination là gold standard.'),
    ('Walkup 2008 CAMS — CBT đơn trị vs Placebo',
     '59,7 % vs 23,7 % → NNT ≈ 3',
     'RẤT TỐT. Ngay cả CBT đơn trị cũng có NNT tương đương combination.'),
    ('QT029 Li 2025 BMC NMA — ACT vs Chứng',
     'MD = −3,83 (CrI chứa 0)',
     'Không thể tính NNT vì chỉ có điểm thang liên tục (không phải response rate). Đây là hạn chế của thang đo liên tục.'),
    ('QT040 Walder 2025 DMHI SAD — DMHI specific vs Waitlist',
     'Hedges g = 0,878 (effect size)',
     'Không tính được NNT từ g. Phải có tỷ lệ nhị phân (response/non-response) mới tính được.'),
]
for i, (k, v, n) in enumerate(projs):
    row = proj_tbl.rows[i+1]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[2].text = n
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True

d.add_paragraph()

add_para('Lưu ý phương pháp luận: NNT chỉ tính được khi có TỶ LỆ NHỊ PHÂN (đáp ứng/không đáp ứng, remission/không remission). Với thang đo LIÊN TỤC (điểm GAD-7, DASS-21, giảm X điểm), phải dùng effect size (Cohen d, Hedges g, SMD). Hai hệ thống này bổ sung cho nhau, không thay thế nhau.', bold=False)

d.add_paragraph()

# =========================================================
# PHẦN 6 — TÓM GỌN
# =========================================================
add_heading('6. Ghi nhớ 3 ý cốt lõi', level=1)

core_tbl = d.add_table(rows=3, cols=1)
core_tbl.style = 'Table Grid'
cores = [
    ('① NNT = 1 / ARR',
     'ARR = chênh tỷ lệ đáp ứng giữa can thiệp và đối chứng. ARR càng lớn → NNT càng nhỏ. Làm tròn NNT LÊN (không xuống).',
     'E2EFDA'),
    ('② Càng nhỏ càng tốt — logic toán',
     'NNT nhỏ = cần ÍT bệnh nhân điều trị = hiệu suất lớn. NNT = 3 nghĩa "chỉ cần 3 người → 1 khỏi thêm", hiệu quả hơn NNT = 30 ("cần 30 người → 1 khỏi thêm").',
     'FFF2CC'),
    ('③ Luôn đọc NNT trong ngữ cảnh',
     'NNT cần gắn với: thời gian theo dõi + định nghĩa response + nhóm đối chứng. Không có "NNT tốt" chung — chỉ có "NNT = X trong hoàn cảnh Y".',
     'DEEBF7'),
]
for i, (k, v, clr) in enumerate(cores):
    cell = core_tbl.rows[i].cells[0]
    shade_cell(cell, clr)
    p = cell.paragraphs[0]
    r1 = p.add_run(k + '\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)

d.add_paragraph()

# TLTK
add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Laupacis A, Sackett DL, Roberts RS. (1988). An assessment of clinically useful measures of the consequences of treatment. New England Journal of Medicine, 318(26):1728-1733. [Bài đầu tiên giới thiệu NNT]',
    'Cook RJ, Sackett DL. (1995). The number needed to treat: a clinically useful measure of treatment effect. BMJ, 310(6977):452-454. [Tổng quan thực hành]',
    'Walkup JT, Albano AM, Piacentini J, et al. (2008). Cognitive behavioral therapy, sertraline, or a combination in childhood anxiety. NEJM, 359:2753-2766. [CAMS — nguồn NNT = 3]',
    'Zugman A, Grillon C, Pine DS. (2024). Treatment of pediatric anxiety disorders. American Journal of Psychiatry. [QT028 — trích NNT từ CAMS]',
    'Citrome L, Ketter TA. (2013). When does a difference make a difference? Interpretation of number needed to treat, number needed to harm, and likelihood to be helped or harmed. International Journal of Clinical Practice, 67(5):407-411. [NNT trong lâm sàng]',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 20/04/2026 | Các response rate 80,7% / 59,7% / 23,7% trích từ QT028 Zugman 2024 về CAMS. Công thức NNT = 1/ARR là chuẩn trong evidence-based medicine (Laupacis 1988, Cook & Sackett 1995).')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/NNT_dien_giai_va_vi_sao_cang_nho_cang_hieu_qua.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
