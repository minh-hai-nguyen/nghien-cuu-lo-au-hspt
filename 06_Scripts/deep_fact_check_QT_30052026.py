# -*- coding: utf-8 -*-
"""Deep fact-check QT summaries vs PDF goc.
- Extract ALL numerical facts (sample size, %, p, OR, M, SD, β, F) tu QT
- Search EACH fact in PDF goc
- Flag mismatch (potential fabrication or wrong PDF labeling)
30/05/2026."""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import fitz

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
QT_FOLDER = os.path.join(ROOT, 'Tom-tat-tung-bai')

# Build QT → PDF mapping from filename
def find_pdf_for_qt(qt_num):
    """Find PDF in 02_Papers-goc matching QT number."""
    for root, dirs, files in os.walk(os.path.join(ROOT, '02_Papers-goc')):
        for f in files:
            if f.startswith(f'QT{qt_num}_') and f.lower().endswith('.pdf'):
                return os.path.join(root, f)
            # Also try without QT prefix (older naming)
            if f.startswith(f'{int(qt_num):02d}_') and f.lower().endswith('.pdf'):
                return os.path.join(root, f)
    return None


def extract_facts_from_text(text):
    """Extract numerical facts with categorization."""
    facts = []
    # Percentages
    for m in re.finditer(r'(\d+[\.,]\d+\s*%|\d+\s*%)', text):
        facts.append(('%', m.group(1).strip(), m.start()))
    # p-values
    for m in re.finditer(r'(p\s*[<>=]\s*0[\.,]\d+|p-?value\s*[<>=]\s*0[\.,]\d+)', text):
        facts.append(('p', m.group(1), m.start()))
    # OR / RR
    for m in re.finditer(r'\b(?:OR|RR|HR)\s*=?\s*(\d+[\.,]\d+)', text):
        facts.append(('OR/RR', m.group(0), m.start()))
    # β / beta
    for m in re.finditer(r'(?:β|beta)\s*=?\s*[+-]?\d+[\.,]\d+', text):
        facts.append(('β', m.group(0), m.start()))
    # F-statistic
    for m in re.finditer(r'F\s*\(?\d*[\.,]?\d*\)?\s*=\s*\d+[\.,]\d+', text):
        facts.append(('F', m.group(0), m.start()))
    # M ± SD or M (SD)
    for m in re.finditer(r'\b\d+[\.,]\d+\s*±\s*\d+[\.,]\d+', text):
        facts.append(('M±SD', m.group(0), m.start()))
    # Sample sizes (large numbers near "students/HS/adolescent/participant")
    for m in re.finditer(r'\b([\d.,]+)\s*(?:học sinh|HS\b|adolescent|student|participant|trẻ)',
                         text, re.IGNORECASE):
        n = m.group(1)
        # Normalize and check size
        num = re.sub(r'[.,]', '', n)
        if num.isdigit() and 50 <= int(num) <= 1000000:
            facts.append(('N', n, m.start()))
    return facts


def normalize_number(s):
    """Normalize Vietnamese (1.234,5) and English (1,234.5) to compare."""
    # Remove all punctuation, get digits + final decimal
    s = s.strip().replace(' ', '')
    # Remove % sign
    has_pct = '%' in s
    s = s.replace('%', '')
    # Try parsing
    # Vietnamese: comma is decimal, period is thousands
    # English: period is decimal, comma is thousands
    # If both present, determine by which comes LAST
    if ',' in s and '.' in s:
        if s.rindex(',') > s.rindex('.'):
            # Vietnamese format
            s = s.replace('.', '').replace(',', '.')
        else:
            s = s.replace(',', '')
    elif ',' in s:
        # Could be decimal or thousands
        parts = s.split(',')
        if len(parts) == 2 and len(parts[1]) <= 2:
            s = parts[0] + '.' + parts[1]  # decimal
        else:
            s = s.replace(',', '')  # thousands
    elif '.' in s:
        parts = s.split('.')
        if len(parts) == 2 and len(parts[1]) >= 3:
            s = s.replace('.', '')  # thousands
    suffix = '%' if has_pct else ''
    return s + suffix


def check_fact_in_pdf(fact_str, pdf_text):
    """Check if fact appears in PDF text with multiple normalizations."""
    # Try direct
    if fact_str in pdf_text:
        return True
    # Try normalized
    fact_norm = normalize_number(fact_str)
    pdf_norm = re.sub(r'[ \t]', '', pdf_text)
    if fact_norm in pdf_norm:
        return True
    # Try swap comma/period
    swap = fact_str.replace(',', 'TEMP').replace('.', ',').replace('TEMP', '.')
    if swap in pdf_text:
        return True
    return False


# Process QT001-QT020 deep check
results = []
qt_files = sorted([f for f in os.listdir(QT_FOLDER)
                   if re.match(r'QT0[01]\d_', f) and f.endswith('.docx')
                   and 'FIXED' not in f and 'BACKUP' not in f])[:20]

print(f'Deep-checking {len(qt_files)} QT files...')
print()

for qt_file in qt_files:
    m = re.match(r'QT(\d{3})_', qt_file)
    qt_num = m.group(1)
    qt_path = os.path.join(QT_FOLDER, qt_file)

    pdf_path = find_pdf_for_qt(qt_num)
    if not pdf_path:
        results.append({
            'qt': qt_num, 'file': qt_file,
            'pdf': None, 'status': 'NO_PDF',
        })
        continue

    # Read QT
    try:
        d = Document(qt_path)
        qt_text = '\n'.join(p.text for p in d.paragraphs)
        for tb in d.tables:
            for row in tb.rows:
                for c in row.cells:
                    qt_text += '\n' + c.text
    except Exception as e:
        results.append({
            'qt': qt_num, 'file': qt_file,
            'pdf': pdf_path, 'status': f'QT_READ_ERROR: {e}',
        })
        continue

    # Read PDF (all pages)
    try:
        doc = fitz.open(pdf_path)
        pdf_text = ''
        for i in range(doc.page_count):
            pdf_text += doc[i].get_text() + '\n'
        doc.close()
    except Exception as e:
        results.append({
            'qt': qt_num, 'file': qt_file,
            'pdf': pdf_path, 'status': f'PDF_READ_ERROR: {e}',
        })
        continue

    # Extract facts from QT
    qt_facts = extract_facts_from_text(qt_text)

    # Check each fact
    matched = 0
    missed = []
    for kind, val, pos in qt_facts:
        if check_fact_in_pdf(val, pdf_text):
            matched += 1
        else:
            missed.append((kind, val))

    total = len(qt_facts)
    rate = matched / total if total > 0 else 0

    results.append({
        'qt': qt_num,
        'file': qt_file,
        'pdf': os.path.basename(pdf_path),
        'total_facts': total,
        'matched': matched,
        'missed': missed[:15],  # top 15 missed
        'match_rate': round(rate * 100, 1),
        'status': 'OK' if rate > 0.7 else 'SUSPICIOUS',
    })


# Print summary
print('=== RESULTS ===')
for r in results:
    if r['status'] == 'NO_PDF':
        print(f'[QT{r["qt"]}] NO PDF available')
        continue
    if 'ERROR' in r['status']:
        print(f'[QT{r["qt"]}] {r["status"]}')
        continue
    print(f'[QT{r["qt"]}] {r["match_rate"]}% match ({r["matched"]}/{r["total_facts"]}) | {r["status"]}')
    if r['status'] == 'SUSPICIOUS' and r['missed']:
        print(f'   MISSED facts (top 10):')
        for kind, val in r['missed'][:10]:
            print(f'     [{kind}] {val}')

# Save
with open(os.path.join(ROOT, '06_Scripts', 'QT_deep_factcheck_30052026.json'),
          'w', encoding='utf-8') as f:
    json.dump(results, f, ensure_ascii=False, indent=2)
print()
print('Saved: 06_Scripts/QT_deep_factcheck_30052026.json')
