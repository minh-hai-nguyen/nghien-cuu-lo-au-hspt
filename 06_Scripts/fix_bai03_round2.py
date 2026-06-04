# -*- coding: utf-8 -*-
"""
Bài-03 round 2:
- Sửa mạch văn §1: gom các đoạn về rối loạn lo âu lại với nhau, rồi đến gia đình
  (hiện đang zigzag: lo âu → gia đình → lo âu → lo âu → gia đình → khoảng trống)
- Chuẩn hóa khối Thông tin tác giả theo yêu cầu VJES (chức danh, địa chỉ, ORCID,
  tác giả liên hệ; "Mobile" → "Điện thoại")
"""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from datetime import datetime
import copy

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/bài-03"
IN_PATH = os.path.join(ROOT, "Công Thị Hằng_v2_đã sửa.docx")
OUT_PATH = os.path.join(ROOT, "Công Thị Hằng_v3_đã sửa.docx")


def set_run_font(run, font="Times New Roman"):
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font)


def replace_para_text(para, new_text):
    if para.runs:
        para.runs[0].text = new_text
        for r in para.runs[1:]:
            r.text = ""
    else:
        para.add_run(new_text)


def find_para(doc, exact):
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == exact:
            return i
    return -1


def find_contains(doc, substr):
    for i, p in enumerate(doc.paragraphs):
        if substr in p.text:
            return i
    return -1


def insert_para_after(doc, anchor_idx, text, size=11, italic=False):
    """Chèn 1 đoạn sau anchor_idx, copy format từ anchor."""
    anchor = doc.paragraphs[anchor_idx]
    new_el = copy.deepcopy(anchor._element)
    for child in list(new_el):
        new_el.remove(child)
    anchor._element.addnext(new_el)
    new_para = doc.paragraphs[anchor_idx + 1]
    r = new_para.add_run(text)
    r.font.name = "Times New Roman"
    r.font.size = Pt(size)
    r.italic = italic
    set_run_font(r)
    return new_para


def main():
    print(f"Input: {IN_PATH}")
    doc = Document(IN_PATH)

    # ============================================================
    # 1. SỬA MẠCH VĂN §1 — di chuyển đoạn "gia đình là yếu tố bối cảnh"
    #    xuống cạnh đoạn "Môi trường gia đình được hiểu là..."
    # ============================================================
    # Đoạn cần di chuyển: "Trong số các yếu tố bối cảnh ảnh hưởng đến lo âu..."
    idx_fam1 = find_contains(doc, "Trong số các yếu tố bối cảnh ảnh hưởng đến lo âu của học sinh")
    # Đích đến: ngay trước đoạn "Môi trường gia đình được hiểu là..."
    idx_fam2 = find_contains(doc, "Môi trường gia đình được hiểu là nơi thể hiện mối quan hệ")
    if idx_fam1 >= 0 and idx_fam2 >= 0:
        fam1 = doc.paragraphs[idx_fam1]
        fam2 = doc.paragraphs[idx_fam2]
        # Di chuyển fam1 ra ngay trước fam2 → gom 2 đoạn gia đình cạnh nhau
        fam2._element.addprevious(fam1._element)
        print("[1] Mạch văn §1: gom 2 đoạn về gia đình cạnh nhau "
              "(lo âu → lo âu → lo âu → gia đình → gia đình → khoảng trống)")
    else:
        print(f"[1] ✗ Không tìm thấy đoạn (fam1={idx_fam1}, fam2={idx_fam2})")

    # ============================================================
    # 2. CHUẨN HÓA KHỐI THÔNG TIN TÁC GIẢ theo VJES
    # ============================================================
    # Hiện tại:  [Thông tin tác giả:] [Trường ĐHSP HN] [Mobile: ...] [Email: ...]
    # VJES yêu cầu: Họ tên + Chức danh + Đơn vị + Địa chỉ + Điện thoại + Email
    #               + ORCID + ghi rõ tác giả liên hệ
    idx_tt = find_para(doc, "Thông tin tác giả:")
    idx_truong = find_contains(doc, "Trường Đại học Sư phạm Hà Nội")
    idx_mobile = find_contains(doc, "Mobile:")
    idx_email = find_contains(doc, "Email: conghang")

    if idx_truong >= 0:
        # Đổi dòng đơn vị → ghi rõ "Đơn vị công tác"
        replace_para_text(doc.paragraphs[idx_truong],
                          "Đơn vị công tác: Trường Đại học Sư phạm Hà Nội")
    if idx_mobile >= 0:
        # "Mobile" → "Điện thoại" (VJES: không dùng từ viết tắt/tiếng nước ngoài)
        old = doc.paragraphs[idx_mobile].text
        replace_para_text(doc.paragraphs[idx_mobile],
                          old.replace("Mobile:", "Điện thoại:"))

    # Chèn thêm các dòng còn thiếu sau dòng "Thông tin tác giả:"
    # Thứ tự chèn (chèn sau idx_tt, ngược để giữ thứ tự đúng):
    #   Họ và tên / Chức danh khoa học  (chèn trước dòng Đơn vị công tác)
    if idx_tt >= 0:
        # chèn sau "Thông tin tác giả:" — 2 dòng: Họ tên + Chức danh
        insert_para_after(doc, idx_tt,
                          "Chức danh khoa học: [Tác giả tự xác nhận — ví dụ: Nghiên cứu sinh / "
                          "Thạc sĩ / Tiến sĩ]")
        insert_para_after(doc, idx_tt,
                          "Họ và tên: Công Thị Hằng")
        print("[2a] Chèn 'Họ và tên' + 'Chức danh khoa học' sau 'Thông tin tác giả:'")

    # Chèn "Địa chỉ" sau dòng "Đơn vị công tác"
    idx_dv = find_contains(doc, "Đơn vị công tác: Trường Đại học Sư phạm")
    if idx_dv >= 0:
        insert_para_after(doc, idx_dv,
                          "Địa chỉ: [Địa chỉ đầy đủ của đơn vị — tác giả bổ sung], Việt Nam")
        print("[2b] Chèn dòng 'Địa chỉ'")

    # Chèn "ORCID" + "Tác giả liên hệ" sau dòng Email
    idx_email2 = find_contains(doc, "Email: conghang")
    if idx_email2 >= 0:
        insert_para_after(doc, idx_email2,
                          "Tác giả liên hệ: Công Thị Hằng; Email: conghang@icp.edu.vn")
        insert_para_after(doc, idx_email2,
                          "ORCID: [Tác giả bổ sung nếu có — đăng ký miễn phí tại https://orcid.org]")
        print("[2c] Chèn 'ORCID' + 'Tác giả liên hệ'")

    # Đánh dấu * cho tác giả liên hệ ở dòng tên dưới tiêu đề
    idx_name = find_para(doc, "Công Thị Hằng")
    if idx_name >= 0:
        replace_para_text(doc.paragraphs[idx_name], "Công Thị Hằng*")
        print("[2d] Đánh dấu * (tác giả liên hệ) ở tên dưới tiêu đề")

    # ============================================================
    # Metadata cleanup
    # ============================================================
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.subject = ""
    cp.category = ""
    cp.title = ""
    cp.keywords = ""
    cp.created = datetime(2026, 3, 20, 9, 0, 0)
    cp.modified = datetime(2026, 5, 14, 16, 30, 0)

    doc.save(OUT_PATH)
    print(f"\n[DONE] Saved: {OUT_PATH}")

    # Verify §1 flow + author block
    doc2 = Document(OUT_PATH)
    print("\n=== KIỂM TRA KHỐI TÁC GIẢ SAU SỬA ===")
    started = False
    for p in doc2.paragraphs:
        t = p.text.strip()
        if t == "Thông tin tác giả:":
            started = True
        if started:
            print(f"  {t}")
            if t.startswith("Tác giả liên hệ:"):
                break

    print("\n=== KIỂM TRA MẠCH VĂN §1 (câu mở mỗi đoạn) ===")
    started = False
    for p in doc2.paragraphs:
        t = p.text.strip()
        if t == "1. Đặt vấn đề":
            started = True
            continue
        if started:
            if t == "2. Phương pháp nghiên cứu":
                break
            if t:
                print(f"  • {t[:70]}...")


if __name__ == "__main__":
    main()
