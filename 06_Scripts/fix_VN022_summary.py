# -*- coding: utf-8 -*-
"""Sua loi tom tat VN022: 249 GV -> 66 GV; Kon Tum/Ninh Thuan -> Gia Lai"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SUM_PATH = os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN022_UNICEF_VN_2022_SchoolFactors.docx')

d = Document(SUM_PATH)
changes = 0
for p in d.paragraphs:
    if p.text.strip():
        new_text = p.text
        # Fix 249 GV -> 66 GV
        new_text = new_text.replace('249 GV', '66 GV')
        new_text = new_text.replace('249 giáo viên', '66 giáo viên')
        # Fix provinces
        new_text = new_text.replace('Hà Nội, Điện Biên, Đồng Tháp, Kon Tum, Ninh Thuận',
                                    'Hà Nội, Điện Biên, Đồng Tháp, Gia Lai (+ TPHCM theo kế hoạch nhưng bỏ do COVID-19)')
        new_text = new_text.replace('Điện Biên, Kon Tum', 'Điện Biên, Gia Lai')
        new_text = new_text.replace('ven biển (Ninh Thuận)',
                                    'Tây Nguyên (Gia Lai)')
        # Fix 47% stat (actual: UNICEF reports 50%+ student study >2h, 28% >3h/night)
        new_text = new_text.replace('47% học thêm >3h/tuần',
                                    '50%+ HS học > 2 giờ/ngày, 28% HS học > 3 giờ/đêm; 50% HS có > 3h học thêm/tuần (15% > 9h)')
        if new_text != p.text:
            # Update paragraph while preserving formatting
            for run in p.runs:
                run.text = ''
            p.runs[0].text = new_text if p.runs else p.add_run(new_text)
            changes += 1

# Also fix in tables
for t in d.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if p.text.strip():
                    nt = p.text
                    nt = nt.replace('47% học thêm >3h/tuần; 15% >9h/tuần',
                                    '50% HS học thêm >3h/tuần; 15% >9h/tuần')
                    if nt != p.text:
                        for run in p.runs:
                            run.text = ''
                        if p.runs:
                            p.runs[0].text = nt
                        changes += 1

# Save as new version
OUT = os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN022_UNICEF_VN_2022_SchoolFactors.docx')
d.save(OUT)
print(f'Saved: {os.path.basename(OUT)}')
print(f'Changes: {changes}')
