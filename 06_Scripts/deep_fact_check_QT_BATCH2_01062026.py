# -*- coding: utf-8 -*-
"""Deep fact-check QT021-QT073 (53 files) - batch 2.
Improved normalize: handle decimal comma/period both ways +
space variants (e.g., "13.9 %" vs "13.9%").
01/06/2026."""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import fitz

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
QT_FOLDER = os.path.join(ROOT, 'Tom-tat-tung-bai')


def find_pdf_for_qt(qt_num):
    for root, dirs, files in os.walk(os.path.join(ROOT, '02_Papers-goc')):
        for f in files:
            if f.startswith(f'QT{qt_num}_') and f.lower().endswith('.pdf'):
                return os.path.join(root, f)
            if f.startswith(f'{int(qt_num):02d}_') and f.lower().endswith('.pdf'):
                return os.path.join(root, f)
    return None


def normalize_for_compare(s):
    """Improved normalization: handle multiple format variants."""
    variants = set()
    variants.add(s.strip())
    # Without spaces
    variants.add(s.replace(' ', ''))
    # Comma → period for decimal
    swap1 = re.sub(r'(\d),(\d)', r'\1.\2', s)
    variants.add(swap1)
    variants.add(swap1.replace(' ', ''))
    # Period → comma for decimal
    swap2 = re.sub(r'(\d)\.(\d)', r'\1,\2', s)
    variants.add(swap2)
    variants.add(swap2.replace(' ', ''))
    return variants


def extract_facts(text):
    facts = []
    for m in re.finditer(r'(\d+[\.,]\d+\s*%|\d+\s*%)', text):
        facts.append(('%', m.group(1).strip()))
    for m in re.finditer(r'(p\s*[<>=]\s*0[\.,]\d+)', text):
        facts.append(('p', m.group(1)))
    for m in re.finditer(r'\b(?:OR|RR|HR)\s*=?\s*(\d+[\.,]\d+)', text):
        facts.append(('OR/RR', m.group(0)))
    for m in re.finditer(r'(?:β|beta)\s*=?\s*[+-]?\d+[\.,]\d+', text):
        facts.append(('β', m.group(0)))
    for m in re.finditer(r'F\s*\(?\d*[\.,]?\d*\)?\s*=\s*\d+[\.,]\d+', text):
        facts.append(('F', m.group(0)))
    for m in re.finditer(r'\b\d+[\.,]\d+\s*±\s*\d+[\.,]\d+', text):
        facts.append(('M±SD', m.group(0)))
    for m in re.finditer(r'\b([\d.,]+)\s*(?:học sinh|HS\b|adolescent|student|participant)',
                         text, re.IGNORECASE):
        n = m.group(1)
        num = re.sub(r'[.,]', '', n)
        if num.isdigit() and 50 <= int(num) <= 1000000:
            facts.append(('N', n))
    return facts


def check_fact(val, pdf_text):
    """Check with multiple variants."""
    pdf_compact = re.sub(r'[ \t]', '', pdf_text)
    for v in normalize_for_compare(val):
        if v in pdf_text or v in pdf_compact:
            return True
    return False


# Process QT021-073
qt_files = sorted([f for f in os.listdir(QT_FOLDER)
                   if re.match(r'QT0[2-7]\d_', f) and f.endswith('.docx')
                   and 'FIXED' not in f and 'BACKUP' not in f])

print(f'Batch 2: {len(qt_files)} files (QT021-QT073)')

results = []
for qt_file in qt_files:
    m = re.match(r'QT(\d{3})_', qt_file)
    qt_num = m.group(1)
    qt_path = os.path.join(QT_FOLDER, qt_file)
    pdf_path = find_pdf_for_qt(qt_num)
    if not pdf_path:
        results.append({'qt': qt_num, 'status': 'NO_PDF', 'file': qt_file})
        continue
    try:
        dq = Document(qt_path)
        qt_text = '\n'.join(p.text for p in dq.paragraphs)
        for tb in dq.tables:
            for row in tb.rows:
                for c in row.cells:
                    qt_text += '\n' + c.text
    except Exception as e:
        results.append({'qt': qt_num, 'status': f'QT_ERR: {e}', 'file': qt_file})
        continue
    try:
        doc = fitz.open(pdf_path)
        pdf_text = ''
        for i in range(doc.page_count):
            pdf_text += doc[i].get_text() + '\n'
        doc.close()
    except Exception as e:
        results.append({'qt': qt_num, 'status': f'PDF_ERR: {e}', 'file': qt_file})
        continue

    qt_facts = extract_facts(qt_text)
    matched = 0
    missed = []
    for kind, val in qt_facts:
        if check_fact(val, pdf_text):
            matched += 1
        else:
            missed.append((kind, val))
    total = len(qt_facts)
    rate = matched / total if total > 0 else 0
    results.append({
        'qt': qt_num,
        'file': qt_file,
        'pdf': os.path.basename(pdf_path),
        'total': total,
        'matched': matched,
        'rate': round(rate * 100, 1),
        'status': 'OK' if rate >= 0.7 else ('LOW' if rate >= 0.3 else 'SUSPICIOUS'),
        'missed_top': missed[:10],
    })

# Print summary
print()
ok_count = sum(1 for r in results if r.get('status') == 'OK')
low_count = sum(1 for r in results if r.get('status') == 'LOW')
sus_count = sum(1 for r in results if r.get('status') == 'SUSPICIOUS')
no_pdf = sum(1 for r in results if r.get('status') == 'NO_PDF')
err = sum(1 for r in results if 'ERR' in str(r.get('status', '')))

print(f'OK (>=70%): {ok_count}')
print(f'LOW (30-70%): {low_count}')
print(f'SUSPICIOUS (<30%): {sus_count}')
print(f'NO_PDF: {no_pdf}')
print(f'ERR: {err}')
print()

# Detail per QT
for r in results:
    if r.get('status') == 'OK':
        print(f'[QT{r["qt"]}] {r["rate"]}% OK')
    elif 'rate' in r:
        st = r['status']
        print(f'[QT{r["qt"]}] {r["rate"]}% {st} | {r["file"][:60]}')
        if r['missed_top']:
            for kind, val in r['missed_top'][:5]:
                print(f'    miss [{kind}] {val}')
    else:
        print(f'[QT{r["qt"]}] {r["status"]} | {r["file"][:60]}')

with open(os.path.join(ROOT, '06_Scripts', 'QT_batch2_factcheck_01062026.json'),
          'w', encoding='utf-8') as f:
    json.dump(results, f, ensure_ascii=False, indent=2)
print()
print('Saved: 06_Scripts/QT_batch2_factcheck_01062026.json')
