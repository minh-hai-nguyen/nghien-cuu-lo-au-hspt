# -*- coding: utf-8 -*-
"""Minor polish for VN022 translation — terminology consistency"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

d = Document(PATH)
changes = 0

# Fix "tuổi teen" → "vị thành niên" (consistency with 314 other uses)
FIXES = [
    ('tuổi teen', 'vị thành niên'),  # informal → formal consistency
]

for p in d.paragraphs:
    if p.text.strip():
        new_text = p.text
        for find, rep in FIXES:
            if find in new_text:
                new_text = new_text.replace(find, rep)
        if new_text != p.text:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_text
            changes += 1

for t in d.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if p.text.strip():
                    new_text = p.text
                    for find, rep in FIXES:
                        if find in new_text:
                            new_text = new_text.replace(find, rep)
                    if new_text != p.text:
                        for run in p.runs:
                            run.text = ''
                        if p.runs:
                            p.runs[0].text = new_text
                        changes += 1

d.save(PATH)
print(f'Minor polish: {changes} changes')
