# -*- coding: utf-8 -*-
"""Áp dụng phong cách CTH v5 vào phần ĐÁNH GIÁ NHANH + PHẢN BIỆN
của các bản dịch 14-29. Chỉ sửa phần mình viết (không sửa dịch gốc)."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

def find_paragraph_index(doc, keyword):
    """Find first paragraph containing keyword"""
    for i, p in enumerate(doc.paragraphs):
        if keyword in p.text:
            return i
    return -1

def insert_after(doc, idx, text, color=None, bold=False, size=12):
    """Insert paragraph after given index"""
    p = doc.paragraphs[idx]
    new_p = OxmlElement('w:p')
    p._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    new_para = Paragraph(new_p, p._element.getparent())
    r = new_para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = color
    return new_para

# CTH v5 style improvements for each file
# Only modify files that have full translations (14, 15, 16)
# Others keep as-is since they're short reviews

files_to_improve = {
    '14_HoangTrungHoc_2025_AJPR.docx': {
        'kt1_insert': 'Nói cách khác, đây là nghiên cứu cắt ngang tại hai thời điểm — không phải nghiên cứu dọc thực sự — so sánh mức độ rối loạn cảm xúc ở thanh thiếu niên Việt Nam trước và sau khi COVID-19 kết thúc.',
        'kt3_gender': 'Bài báo ghi nhận giới tính có ảnh hưởng (Beta = 0,053, P = 0,001) nhưng không phân tích riêng tỷ lệ lo âu nam vs nữ — đây là một thiếu sót phương pháp luận khi hầu hết nghiên cứu quốc tế đều so sánh giới chi tiết (McLean và cộng sự, 2011).',
        'arch_b': 'Dữ liệu của nghiên cứu này, cho thấy sự giảm đáng kể các triệu chứng cảm xúc tiêu cực sau đại dịch (lo âu giảm từ 41,5% xuống 25,4%), gợi ý rằng thanh thiếu niên có khả năng phục hồi nhất định nhưng tỷ lệ hậu COVID vẫn cao hơn trước dịch — điều này đòi hỏi tiếp tục hỗ trợ tâm lý dài hạn.',
    },
    '15_NgoAnhVinh_2024_JAffectDisordReports.docx': {
        'kt1_insert': 'Nói cách khác, nghiên cứu này kết hợp hai chiều đánh giá — tình trạng sức khỏe tâm thần hiện tại (DASS-21) và trải nghiệm bất lợi quá khứ (ACEs) — để hiểu mối quan hệ giữa "vết thương thời thơ ấu" và rối loạn cảm xúc ở thanh thiếu niên dân tộc thiểu số.',
        'kt3_gender': 'Tổng quan tài liệu cho thấy nữ giới thường có tỷ lệ lo âu cao hơn nam (Salk và cộng sự, 2017). Tuy nhiên, nghiên cứu này không phân tích giới tính chi tiết cho lo âu — chỉ có OR = 1,29 cho căng thẳng ở nữ, không đạt ý nghĩa (P > 0,05). Đây là khoảng trống cần lấp đầy trong bối cảnh văn hóa giới ở vùng dân tộc thiểu số có thể khác biệt.',
        'arch_b': 'Dữ liệu của nghiên cứu này, cho thấy gần 60% thanh thiếu niên dân tộc thiểu số nội trú có triệu chứng trầm cảm và gần 50% đã trải qua ít nhất một trải nghiệm bất lợi, gợi ý rằng nhóm này cần được ưu tiên trong chính sách can thiệp sức khỏe tâm thần quốc gia — đặc biệt thông qua cải thiện mối quan hệ bạn bè tại trường nội trú.',
    },
    '16_BaoQuyen_2025_YHCD.docx': {
        'kt1_insert': 'Nói cách khác, nghiên cứu sử dụng thang sàng lọc DASS-21 với ngưỡng cắt rất thấp (lo âu ≥4) — bao gồm cả mức "nhẹ" mà có thể không cần can thiệp lâm sàng — nên tỷ lệ 86,2% cần được diễn giải thận trọng.',
        'kt3_gender': 'Y văn quốc tế xác nhận nữ > nam về lo âu (McLean, 2011). Nghiên cứu này phù hợp xu hướng đó (p < 0,05). Tuy nhiên, mẫu có 78% nữ — thiên lệch nghiêm trọng — nên kết quả giới tính có thể phản ánh đặc điểm mẫu hơn là dân số thực.',
        'arch_b': 'Dữ liệu của nghiên cứu này, cho thấy tỷ lệ lo âu 86,2% — cao nhất trong tất cả nghiên cứu tại Việt Nam, gợi ý rằng hoặc DASS-21 với ngưỡng thấp không phù hợp cho sàng lọc dân số chung, hoặc sức khỏe tâm thần học sinh THPT Hà Nội thực sự ở mức báo động — cần nghiên cứu xác nhận với công cụ chẩn đoán (DISC-5).',
    },
}

for fname, improvements in files_to_improve.items():
    path = fname
    if not os.path.exists(path):
        sys.stderr.write(f'SKIP {fname} — not found\n')
        continue

    doc = Document(path)

    # Find TÓM TẮT ĐÁNH GIÁ NHANH section and add KT1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip()
        # Add KT1 after "Hạn chế:" line in ĐÁNH GIÁ NHANH
        if text == 'Hạn chế:' and i < 20:
            # Find the last bullet point after Hạn chế
            j = i + 1
            while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip().startswith('\u2022'):
                j += 1
            # Insert KT1 observation
            if 'kt1_insert' in improvements:
                new_p = doc.add_paragraph()
                r = new_p.add_run(improvements['kt1_insert'])
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
                r.italic = True
                r.font.color.rgb = BLUE
                # Move to after hạn chế section
                doc.paragraphs[j-1]._element.addnext(new_p._element)
            break

    # Add Architecture B conclusion near PHẢN BIỆN
    for i, p in enumerate(doc.paragraphs):
        if 'PHẢN BIỆN' in p.text.upper() and p.style.name.startswith('Heading'):
            if 'arch_b' in improvements:
                # Add before PHẢN BIỆN
                new_p = doc.add_paragraph()
                r = new_p.add_run('Nhận định tổng hợp (Kiến trúc B): ')
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE
                r2 = new_p.add_run(improvements['arch_b'])
                r2.font.name = 'Times New Roman'; r2.font.size = Pt(12); r2.font.color.rgb = BLUE
                p._element.addprevious(new_p._element)

            if 'kt3_gender' in improvements:
                # Add gender analysis after PHẢN BIỆN heading
                new_p2 = doc.add_paragraph()
                r3 = new_p2.add_run('Đối chiếu giới tính (KT3): ')
                r3.bold = True; r3.font.name = 'Times New Roman'; r3.font.size = Pt(11); r3.font.color.rgb = RED
                r4 = new_p2.add_run(improvements['kt3_gender'])
                r4.font.name = 'Times New Roman'; r4.font.size = Pt(11); r4.font.color.rgb = RED
                # Insert after PHẢN BIỆN heading
                p._element.addnext(new_p2._element)
            break

    doc.save(fname)
    sys.stderr.write(f'CTH v5 applied: {fname}\n')

sys.stderr.write('\nDone — CTH v5 style applied to 3 full translations\n')
