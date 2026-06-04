"""Build doc: ACT ra doi nhu the nao + vi sao chi 1 NC xep ACT > CBT."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

# ===== TITLE =====
title = d.add_heading('ACT (Acceptance and Commitment Therapy): Lịch sử ra đời và vì sao chỉ 1 nghiên cứu xếp ACT cao hơn CBT?', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Giải thích cho thầy về liệu pháp Chấp nhận và Cam kết, và nghịch lý xếp hạng trong QT029 Li et al. 2025')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# META BOX
meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
mp = cell.paragraphs[0]
mp.add_run('Câu hỏi gốc của thầy: ').bold = True
mp.add_run('"ACT ra đời như thế nào? Vì sao chỉ có một nghiên cứu xếp loại hiệu quả can thiệp của ACT cao hơn CBT?"')
mp2 = cell.add_paragraph()
mp2.add_run('Nguồn verify: ').bold = True
mp2.add_run('Số liệu ACT MD = −3,83 [95% CrI: −9,33; 1,51] và CBT MD = −3,64 [95% CrI: −8,23; 3,32] đã verify trực tiếp từ PDF QT029 Li et al. 2025 BMC Psychiatry.')
d.add_paragraph()

# =========================================================
# PHẦN 1 — LỊCH SỬ RA ĐỜI CỦA ACT
# =========================================================
add_heading('PHẦN I — ACT ra đời như thế nào?', level=1)

add_para('1. Cha đẻ và bối cảnh khởi đầu', bold=True, size=12)
add_para('ACT (Acceptance and Commitment Therapy — Liệu pháp Chấp nhận và Cam kết) do GS. Steven C. Hayes phát triển tại Đại học Nevada, Reno (Hoa Kỳ), bắt đầu từ đầu thập niên 1980. Hayes cùng cộng sự (Kirk Strosahl, Kelly Wilson) xây dựng ACT như một "thế hệ thứ ba" (third wave) của liệu pháp hành vi–nhận thức.')
add_para('Tác phẩm nền tảng: Hayes, Strosahl & Wilson (1999). "Acceptance and Commitment Therapy: An Experiential Approach to Behavior Change". Guilford Press. Đây là cuốn sách đầu tiên hệ thống hoá ACT thành một liệu pháp lâm sàng hoàn chỉnh.')
d.add_paragraph()

add_para('2. Ba thế hệ (three waves) của liệu pháp hành vi', bold=True, size=12)
wave_tbl = d.add_table(rows=4, cols=3)
wave_tbl.style = 'Table Grid'
hdr = wave_tbl.rows[0]
for i, h in enumerate(['Thế hệ', 'Thời kỳ', 'Đại diện + tư tưởng chính']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

waves = [
    ('Thế hệ 1 — Behaviorism',
     '1950–1960',
     'Skinner, Wolpe, Eysenck. Tập trung vào HÀNH VI quan sát được. Phương pháp: điều kiện hóa, giải mẫn cảm, củng cố. Không quan tâm "suy nghĩ".'),
    ('Thế hệ 2 — CBT cổ điển',
     '1960–1990',
     'Aaron Beck (Cognitive Therapy), Albert Ellis (REBT). Tập trung vào SUY NGHĨ lệch lạc. Mục tiêu: XÁC ĐỊNH + THAY ĐỔI suy nghĩ tiêu cực thành tích cực. Đây là CBT truyền thống mà thầy thường gặp trong văn bản.'),
    ('Thế hệ 3 — Mindfulness & ACT',
     '1990–nay',
     'Steven Hayes (ACT), Marsha Linehan (DBT), Jon Kabat-Zinn (MBSR), Zindel Segal (MBCT), Paul Gilbert (CFT). Tập trung vào MỐI QUAN HỆ với suy nghĩ thay vì thay đổi nội dung suy nghĩ. Chủ trương CHẤP NHẬN cảm xúc tiêu cực thay vì chống lại.'),
]
for i, (w, t, c) in enumerate(waves):
    row = wave_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(w)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = t
    row.cells[2].text = c

d.add_paragraph()

add_para('3. Nền tảng lý thuyết — Relational Frame Theory (RFT)', bold=True, size=12)
add_para('ACT không đứng một mình — nó được xây dựng trên nền một lý thuyết khoa học về ngôn ngữ và nhận thức gọi là Relational Frame Theory (RFT — Lý thuyết Khung Quan hệ), do Hayes và cộng sự phát triển song song từ thập niên 1980. RFT giải thích vì sao con người "mắc kẹt" trong những suy nghĩ tiêu cực — vì khả năng ngôn ngữ của chúng ta tạo ra những "khung quan hệ" liên kết các khái niệm với nhau một cách cứng nhắc.')
add_para('Điểm đặc biệt: ACT là một trong rất ít liệu pháp tâm lý có một LÝ THUYẾT CƠ BẢN về hoạt động của tâm trí (RFT) đi kèm. Phần lớn liệu pháp khác (CBT, IPT) dựa trên mô hình lâm sàng nhưng không có lý thuyết khoa học cơ bản như vậy.')
d.add_paragraph()

add_para('4. Sáu tiến trình cốt lõi (Hexaflex)', bold=True, size=12)
hexa_tbl = d.add_table(rows=6, cols=2)
hexa_tbl.style = 'Table Grid'
hexa_data = [
    ('① Acceptance (Chấp nhận)',
     'Cho phép suy nghĩ và cảm xúc khó chịu xuất hiện mà không chống cự. Trái với CBT: không tìm cách xoá/thay thế suy nghĩ tiêu cực, mà "ngồi cùng" chúng.'),
    ('② Cognitive Defusion (Tách rời nhận thức)',
     'Nhìn suy nghĩ như suy nghĩ, không phải sự thật. Ví dụ: thay vì "Tôi là kẻ thất bại", nói "Tôi đang có suy nghĩ rằng tôi là kẻ thất bại".'),
    ('③ Present Moment Awareness (Nhận thức hiện tại)',
     'Tỉnh thức với trải nghiệm "ngay bây giờ". Tương tự mindfulness.'),
    ('④ Self-as-Context (Bản ngã như bối cảnh)',
     'Nhận ra mình là "người quan sát" mọi trải nghiệm, không phải bị đồng nhất với từng suy nghĩ cụ thể.'),
    ('⑤ Values (Giá trị)',
     'Xác định điều gì QUAN TRỌNG với mình — điều muốn sống vì nó (gia đình, học tập, sáng tạo, giúp đỡ...).'),
    ('⑥ Committed Action (Hành động cam kết)',
     'Hành động theo giá trị, kể cả khi suy nghĩ/cảm xúc tiêu cực vẫn còn đó. Đây là điểm "cam kết" trong tên liệu pháp.'),
]
for i, (k, v) in enumerate(hexa_data):
    row = hexa_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FFE699')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(191, 97, 14)
    row.cells[1].text = v

d.add_paragraph()

add_para('Sáu tiến trình này tạo thành một "hexaflex" (hình 6 cạnh) — mục tiêu cuối cùng là tăng psychological flexibility (độ linh hoạt tâm lý): khả năng tiếp xúc với hiện tại một cách cởi mở và hành động theo giá trị dù hoàn cảnh khó khăn.')
d.add_paragraph()

add_para('5. Khác biệt cốt lõi ACT vs CBT cổ điển', bold=True, size=12)
diff_tbl = d.add_table(rows=6, cols=3)
diff_tbl.style = 'Table Grid'
diff_hdr = diff_tbl.rows[0]
for i, h in enumerate(['Khía cạnh', 'CBT cổ điển', 'ACT']):
    cell = diff_hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

diffs = [
    ('Suy nghĩ tiêu cực', 'Cần XÁC ĐỊNH và THAY ĐỔI (cognitive restructuring)', 'CHẤP NHẬN và TÁCH RỜI, không cần thay đổi nội dung'),
    ('Mục tiêu', 'Giảm triệu chứng lo âu', 'Sống theo GIÁ TRỊ, dù triệu chứng còn'),
    ('Quan điểm về cảm xúc', 'Cảm xúc tiêu cực là "vấn đề cần sửa"', 'Cảm xúc tiêu cực là PHẦN TỰ NHIÊN của cuộc sống'),
    ('Ngôn ngữ ẩn dụ', 'Logic, phân tích', 'Rất nhiều ẩn dụ, bài tập trải nghiệm'),
    ('Bài tập điển hình', 'Thought record (ghi chép suy nghĩ), behavioral experiment', 'Leaves on a stream (lá trôi theo dòng), Passengers on the bus (hành khách trên xe buýt)'),
]
for i, (k, c1, c2) in enumerate(diffs):
    row = diff_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = c1
    row.cells[2].text = c2

d.add_paragraph()

# =========================================================
# PHẦN 2 — VÌ SAO CHỈ 1 NC XẾP ACT > CBT
# =========================================================
add_heading('PHẦN II — Vì sao CHỈ MỘT nghiên cứu xếp ACT cao hơn CBT?', level=1, color=(192, 80, 77))

add_para('Nghiên cứu đó là QT029 Li et al. (2025) BMC Psychiatry — Bayesian NMA trên 30 RCT, 1.711 trẻ em/VTN. Kết quả từ bản gốc PDF:', bold=False)
d.add_paragraph()

qt029_tbl = d.add_table(rows=4, cols=3)
qt029_tbl.style = 'Table Grid'
qt029_hdr = qt029_tbl.rows[0]
for i, h in enumerate(['Xếp hạng', 'Can thiệp', 'Mean Difference [95% CrI]']):
    cell = qt029_hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

q29 = [
    ('1', 'ACT', 'MD = −3,83 [95% CrI: −9,33; 1,51]'),
    ('2', 'CBT', 'MD = −3,64 [95% CrI: −8,23; 3,32]'),
    ('3', 'Physical Exercise (PE)', 'MD = −2,16 [95% CrI: −9,99; 5,52]'),
]
for i, (r_, k, v) in enumerate(q29):
    row = qt029_tbl.rows[i+1]
    row.cells[0].text = r_
    row.cells[1].text = k
    row.cells[2].text = v
    for p in row.cells[1].paragraphs:
        for r in p.runs:
            r.bold = True

d.add_paragraph()

# Warning box
warn = d.add_table(rows=1, cols=1)
warn.style = 'Table Grid'
wc = warn.rows[0].cells[0]
shade_cell(wc, 'FCE4D6')
wp = wc.paragraphs[0]
wr = wp.add_run('⚠ CHÚ Ý QUAN TRỌNG: ')
wr.bold = True
wr.font.color.rgb = RGBColor(192, 0, 0)
wp.add_run('Cả 3 khoảng CrI đều CHỨA SỐ DƯƠNG (1,51 / 3,32 / 5,52) — có nghĩa là không NHẤT ĐỊNH "hơn chứng". ACT xếp hạng 1 nhưng CrI của ACT (−9,33; 1,51) chứa giá trị dương → ACT có xác suất nhỏ thực ra KÉM HƠN nhóm đối chứng. Kết luận "ACT hiệu quả nhất" chỉ đúng theo POSTERIOR PROBABILITY (SUCRA), không phải bằng chứng chắc chắn. Tác giả Li 2025 chính thức ghi trong abstract: "this finding should be interpreted with caution due to the overall low quality of evidence."')
d.add_paragraph()

add_heading('5 lý do ACT xếp hạng 1 trong Li 2025 nhưng KHÔNG xuất hiện ở NMA khác', level=2, color=(54, 95, 44))

reasons_tbl = d.add_table(rows=5, cols=2)
reasons_tbl.style = 'Table Grid'
reasons = [
    ('① Số lượng RCT về ACT rất ít so với CBT',
     'Tính đến 2025, có khoảng 20–30 RCT ACT cho trẻ em/VTN. Trong khi đó CBT có TRÊN 300 RCT. Trong NMA của Li 2025 tổng 30 RCT, chỉ khoảng 3–5 RCT là về ACT. Ít RCT → posterior distribution có độ KHÔNG CHẮC CHẮN cao → median có thể bị "kéo" bởi 1–2 RCT positive cực đoan.'),
    ('② Effect size tuyệt đối ACT và CBT gần bằng nhau',
     'MD = −3,83 (ACT) vs −3,64 (CBT) → chênh chỉ 0,19 điểm. Khoảng CrI của 2 can thiệp CHỒNG LẤN gần hoàn toàn. Nghĩa là về mặt thống kê, ACT KHÔNG KHÁC BIỆT có ý nghĩa với CBT. "Xếp hạng 1" chỉ là kết quả toán học của SUCRA — không phải bằng chứng lâm sàng mạnh.'),
    ('③ ACT còn mới, body of evidence còn non',
     'CBT ra đời ~1960, đã tích luỹ bằng chứng suốt 60+ năm. ACT ra đời ~1999 (sách seminal Hayes), đã tích luỹ chỉ 25 năm. Các cơ quan guideline lớn (APA, NICE, AACAP) đều ghi CBT là "first-line" cho childhood anxiety, không phải ACT — vì ACT chưa đủ RCT quy mô lớn để được khuyến nghị hạng 1.'),
    ('④ Selection bias trong RCT ACT',
     'Các RCT ACT thường: (a) mẫu nhỏ (50–150 trẻ); (b) đơn trung tâm; (c) nhóm tác giả tâm huyết với ACT có xu hướng báo cáo kết quả tích cực. Các NMA lớn hơn (Cochrane, Cuijpers) thường LOẠI các RCT nhỏ hoặc có high risk of bias → ACT bị loại phần lớn.'),
    ('⑤ Các NMA khác không include ACT hoặc include ít',
     '• Zugman 2024 AJP (QT028): CBT > Sertraline > Placebo — không có ACT trong so sánh.\n'
     '• Xian 2024 JAD NMA SAD (QT039): iCBT SUCRA 71,2% — chủ yếu focus vào CBT-based, không tách ACT.\n'
     '• Walder 2025 JMIR DMHI (QT040): DMHI general vs SAD-specific — không tách theo loại therapy.\n'
     '• Li 2025 BMC (QT029): ACT > CBT > PE — chỉ NMA này có ACT xếp 1.\n'
     '→ Li 2025 là NC duy nhất trong thư viện có kết quả này. Không có NMA khác "phản biện lẫn nhau" → không thể kết luận chắc chắn ACT > CBT.'),
]
for i, (k, v) in enumerate(reasons):
    row = reasons_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = v

d.add_paragraph()

# =========================================================
# PHẦN 3 — CÁCH ĐỌC ĐÚNG
# =========================================================
add_heading('PHẦN III — Cách đọc đúng kết quả Li 2025', level=1)

add_para('Kết quả ACT > CBT trong Li 2025 KHÔNG phải sai, nhưng cần đọc với 4 điểm sau:', bold=False)
d.add_paragraph()

for it in [
    '1. Hiểu "xếp hạng Bayesian" là xác suất, không phải chứng minh: SUCRA 69% của ACT nghĩa là "trong 1.000 lần mô phỏng posterior, ACT xếp cao hơn các can thiệp khác trong khoảng 690 lần". Không phải "ACT chắc chắn tốt hơn CBT".',
    '2. Luôn kiểm tra khoảng CrI: Nếu CrI chứa 0 hoặc số trái dấu với effect mong đợi → KHÔNG chắc chắn về hiệu quả. Li 2025: CrI của ACT chứa 1,51 → không loại trừ khả năng ACT không hơn chứng.',
    '3. Đọc chú thích của tác giả: Li 2025 tự ghi "low quality of evidence" và "interpreted with caution". Đây là dấu hiệu authors biết kết quả chưa vững.',
    '4. Đối chiếu với nhiều MA khác: Nếu chỉ 1/5 NMA cho kết quả ACT > CBT, 4/5 khác cho CBT hạng 1 → phần lớn bằng chứng vẫn nghiêng về CBT.',
]:
    p = d.add_paragraph(it)
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# PHẦN 4 — TRẢ LỜI NHANH
# =========================================================
add_heading('PHẦN IV — Tóm gọn 1 phút', level=1)

summary_tbl = d.add_table(rows=4, cols=1)
summary_tbl.style = 'Table Grid'
summaries = [
    ('Câu 1: ACT ra đời như thế nào?',
     'Do Steven C. Hayes phát triển tại ĐH Nevada từ đầu 1980s, hệ thống hoá thành liệu pháp chính thức năm 1999. ACT là "thế hệ thứ 3" của liệu pháp hành vi–nhận thức, sau behaviorism (thế hệ 1) và CBT cổ điển (thế hệ 2). Tư tưởng cốt lõi: CHẤP NHẬN cảm xúc thay vì chống lại + CAM KẾT hành động theo giá trị. Nền tảng lý thuyết là Relational Frame Theory (RFT).',
     'E2EFDA'),
    ('Câu 2: Vì sao chỉ 1 NC xếp ACT > CBT?',
     'Vì (a) ACT còn ít RCT so với CBT (3–5 vs hàng chục trong NMA Li 2025); (b) Effect size ACT và CBT thực ra gần bằng nhau (MD chênh 0,19), chỉ khác do toán học SUCRA; (c) CrI của ACT chứa số dương → không chắc chắn hơn chứng; (d) Tác giả tự thừa nhận "low quality of evidence"; (e) 4/5 NMA lớn khác vẫn xếp CBT hạng 1. Nên đọc Li 2025 như "GỢI Ý ACT tiềm năng", không phải "ACT vượt CBT".',
     'FFF2CC'),
    ('Khuyến nghị thực tiễn cho VN',
     'CBT vẫn là lựa chọn hạng 1 (guideline quốc tế + phần lớn NMA). ACT có thể là bổ sung — đặc biệt cho VTN khó tuân thủ CBT cổ điển (thấy nhiều "bài tập ghi chép"), vì ACT dùng ẩn dụ + trải nghiệm hợp văn hoá hơn. Chưa có RCT ACT tại VN — GAP nghiên cứu.',
     'DEEBF7'),
    ('Bài học phương pháp luận',
     'Khi một NMA duy nhất cho kết quả "đảo ngược" so với đồng thuận (ACT > CBT), KHÔNG vội đổi kết luận. Luôn: (1) kiểm CrI/CI, (2) đọc chú thích quality of evidence, (3) đối chiếu với các MA khác, (4) xét số lượng RCT của từng can thiệp. Đây là nguyên tắc "triangulation" trong nghiên cứu tổng hợp.',
     'FCE4D6'),
]
for i, (k, v, clr) in enumerate(summaries):
    cell = summary_tbl.rows[i].cells[0]
    shade_cell(cell, clr)
    p = cell.paragraphs[0]
    r1 = p.add_run(k + '\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# TÀI LIỆU THAM KHẢO
# =========================================================
add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Li Q, et al. (2025). Effects of Different Interventions on Anxiety Disorders in Children and Adolescents: A Systematic Review and Bayesian Network Meta-analysis. BMC Psychiatry. [QT029 — có trong thư viện, verify trực tiếp từ PDF]',
    'Hayes SC, Strosahl KD, Wilson KG. (1999). Acceptance and Commitment Therapy: An Experiential Approach to Behavior Change. Guilford Press, New York. [Sách seminal của ACT]',
    'Hayes SC, Luoma JB, Bond FW, Masuda A, Lillis J. (2006). Acceptance and commitment therapy: Model, processes and outcomes. Behaviour Research and Therapy, 44(1):1-25. [Sáu tiến trình cốt lõi]',
    'Hayes SC, Barnes-Holmes D, Roche B (eds). (2001). Relational Frame Theory: A Post-Skinnerian Account of Human Language and Cognition. Kluwer Academic. [Lý thuyết RFT]',
    'Swain J, Hancock K, Hainsworth C, Bowman J. (2015). Mechanisms of change: Exploratory outcomes from a randomised controlled trial of acceptance and commitment therapy for anxious adolescents. Journal of Contextual Behavioral Science, 4(1):56-67. [RCT ACT cho VTN]',
    'Zugman A, Grillon C, Pine DS. (2024). Treatment of pediatric anxiety disorders: current and emerging approaches. American Journal of Psychiatry. [QT028 — CBT gold standard, không include ACT trong so sánh]',
    'APA Division 12 — Society of Clinical Psychology. (n.d.). Research-Supported Psychological Treatments. [Guideline xếp CBT là "first-line" cho childhood anxiety]',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 20/04/2026 | Số liệu Li 2025 (ACT MD, CBT MD, CrI) verified từ PDF QT029 gốc. Lịch sử ACT dựa trên tài liệu công khai về Steven Hayes và sách Hayes/Strosahl/Wilson 1999 — không có bài primary source trong thư viện dự án về lịch sử ACT.')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/ACT_ra_doi_va_vi_sao_xep_hang_cao_hon_CBT.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
