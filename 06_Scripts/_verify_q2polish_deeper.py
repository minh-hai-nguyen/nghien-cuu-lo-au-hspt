"""Deeper investigation of Q1v7 vs Q2v1 drifts: N=1352, year 2006/2018, bilingual pairs."""
import re
import json
from pathlib import Path
from docx import Document

ROOT = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au")
Q1V7 = ROOT / "bai-bao-Q1" / "Draft_Q1_SongNgu_v7_01062026.docx"
Q2V1 = ROOT / "bai-bao-Q1" / "Draft_Q2_SongNgu_v1_07062026.docx"


def doc_paragraphs(path):
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    # include table cells
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        paras.append("[TABLE] " + p.text)
    return paras


def find_lines(paras, pattern, ctx=0):
    rx = re.compile(pattern)
    hits = []
    for i, p in enumerate(paras):
        if rx.search(p):
            hits.append((i, p))
    return hits


def main():
    out = {}
    for label, path in [("Q1v7", Q1V7), ("Q2v1", Q2V1)]:
        paras = doc_paragraphs(path)
        out[label] = {
            "N1352_lines": [p for _, p in find_lines(paras, r"1[.,]?352|N\s*=\s*1")],
            "year_2006_lines": [p for _, p in find_lines(paras, r"\b2006\b")],
            "year_2018_lines": [p for _, p in find_lines(paras, r"\b2018\b")],
            "Niwenahisemo_lines": [p for _, p in find_lines(paras, r"Niwenahisemo|Hong,\s*S|Su,\s*H")],
            "Anderson_lines": [p for _, p in find_lines(paras, r"Anderson")],
            "Karasu_lines": [p for _, p in find_lines(paras, r"Karasu|3094389")],
            "Rose_lines": [p for _, p in find_lines(paras, r"Rose,?\s*A|co-rumination|đồng tự sự|đồng nghiền ngẫm")],
            "co_rumination_paragraphs": [p for _, p in find_lines(paras, r"co-rumination|đồng (tự sự|nghiền)")],
            "n_paragraphs": len(paras),
        }

    # Bilingual pairs spot check: typical structure is EN para then VN para (or vice versa)
    # Sample first 30 paragraphs of Q2v1 and identify EN vs VN by simple heuristic
    paras_q2 = doc_paragraphs(Q2V1)
    bilingual_samples = []
    for i in range(min(80, len(paras_q2) - 1)):
        p = paras_q2[i]
        nxt = paras_q2[i+1]
        # Look for numbers shared between consecutive paragraphs
        nums_a = set(re.findall(r"\d+[.,]?\d*", p))
        nums_b = set(re.findall(r"\d+[.,]?\d*", nxt))
        shared = nums_a & nums_b
        if len(shared) >= 2 and len(p) > 60 and len(nxt) > 60:
            bilingual_samples.append({
                "i": i,
                "para_a": p[:300],
                "para_b": nxt[:300],
                "shared_numbers": sorted(shared),
            })
    out["bilingual_random_samples"] = bilingual_samples[:8]

    out_path = ROOT / "06_Scripts" / "_verify_q2polish_deeper.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print(f"Wrote {out_path}")


if __name__ == "__main__":
    main()
