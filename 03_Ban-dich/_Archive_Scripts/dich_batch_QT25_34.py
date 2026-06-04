# -*- coding: utf-8 -*-
"""Batch dịch QT25-34 — 10 bài còn lại"""
import sys, os, fitz
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

PDF_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '02_Papers-goc', 'The-gioi-moi')
CHARTS_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), 'Charts')
OUT_DIR = os.path.dirname(os.path.abspath(__file__))

def crop_page(pdf_path, page_idx, y_start_frac, y_end_frac, out_name):
    """Crop a portion of a PDF page as image"""
    d = fitz.open(pdf_path)
    page = d[page_idx]
    rect = fitz.Rect(0, page.rect.height * y_start_frac, page.rect.width, page.rect.height * y_end_frac)
    pix = page.get_pixmap(dpi=200, clip=rect)
    path = os.path.join(CHARTS_DIR, out_name)
    pix.save(path)
    return path

def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    import docx as dx
    d2 = dx.Document(path)
    t = '\n'.join([p.text for p in d2.paragraphs])
    print(f'  -> {name}: {len(t)} chars, {len(d2.tables)} tables')

# =====================================================================
# QT25 — EpiPsychSci 2025 — Crisp et al. — Complete Mental Health Australia
# =====================================================================
def dich_QT25():
    print('QT25 EpiPsychSci...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1017/S2045796025000083', size=10)
    add_heading(doc, 'Sức khỏe tâm thần toàn diện của thanh thiếu niên và người trưởng thành mới nổi tại Úc: Căng thẳng và hạnh phúc qua 3 mẫu cộng đồng đại diện quốc gia', 1)
    h = doc.add_paragraph(); r = h.add_run('The Complete Mental Health of Australia\'s Adolescents and Emerging Adults: Distress and Wellbeing across 3 Nationally Representative Community Samples'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'The Complete Mental Health of Australia\'s Adolescents and Emerging Adults: Distress and Wellbeing across 3 Nationally Representative Community Samples'),
        ('Tác giả', 'Dimity Crisp, Debra Rickwood, Richard Burns, Emily Bariola'),
        ('Cơ quan', 'Đại học Canberra; Đại học Quốc gia Úc; Kantar Public, Melbourne'),
        ('Tạp chí', 'Epidemiology and Psychiatric Sciences (Q1, IF ≈ 7,0)'),
        ('Xuất bản', '2025, Vol. 34, e16, 11 trang'),
        ('DOI', '10.1017/S2045796025000083'),
        ('Loại NC', 'Cắt ngang lặp lại (repeated cross-sectional) — 3 khảo sát quốc gia'),
        ('Mẫu', '2018: n=3.721; 2020: n=974; 2022: n=961 — VTN và thanh niên 12–25 tuổi, Úc'),
    ])
    add_page_ref(doc, '1–11', 'Epidemiology and Psychiatric Sciences', 'Vol. 34, 2025')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Mục tiêu: Mức độ căng thẳng tâm lý cao ở giới trẻ là lo ngại ngày càng tăng. Tuy nhiên, ít khảo sát quốc gia mô tả quỹ đạo SKTT và hạnh phúc qua giai đoạn VTN đến trưởng thành sớm. Nghiên cứu hiện tại tập trung không chỉ vào bệnh tật tâm thần mà còn SKTT tích cực, sử dụng khung Sức khỏe Tâm thần Toàn diện (CMH — Complete Mental Health) của Keyes.')

    p = doc.add_paragraph()
    r = p.add_run('Phương pháp: Người tham gia hoàn thành Khảo sát SKTT Thanh niên Quốc gia năm 2018 (n=3.832), 2020 (n=974) hoặc 2022 (n=961). Áp dụng khung CMH của Keyes để phân loại SKTT và hạnh phúc, kiểm tra tỷ lệ CMH theo thời gian, tuổi và giới. Đo bằng K10 (căng thẳng) và MHC-SF (hạnh phúc).')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Khoảng một nửa báo cáo "phát triển tốt" (flourishing — hạnh phúc cao, không bệnh tật). Tỷ lệ flourishing GIẢM từ 53,0% (2018) xuống 44,4% (2022). Tỷ lệ "loạng choạng" (stumbling — hạnh phúc trung bình, có nguy cơ bệnh) TĂNG từ 19,3% lên 27,3%. Flourishing giảm theo tuổi và phổ biến hơn ở nam so với nữ. Khoảng cách giới MỞ RỘNG: nam flourishing 52,2% vs nữ 36,3% năm 2022.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_p(doc, 'Kết luận: SKTT giới trẻ Úc đang xấu đi — cả về tăng căng thẳng và giảm hạnh phúc. Tiếp cận CMH toàn diện cung cấp bức tranh đầy đủ hơn chỉ đo bệnh tật.')

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['EpiPsychSci Q1 IF≈7 — tạp chí dịch tễ SKTT uy tín.', 'Khung CMH của Keyes — đo CẢ bệnh tật VÀ hạnh phúc, không chỉ pathology.', 'Flourishing giảm 53%→44,4% (2018→2022); stumbling tăng 19,3%→27,3%.', 'Khoảng cách giới mở rộng: nam flourishing 52,2% vs nữ 36,3% (2022).', '3 khảo sát quốc gia đại diện — phương pháp nhất quán.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Chỉ Úc — khác biệt văn hóa với VN/châu Á.', 'Cắt ngang lặp lại — không theo dõi cùng cá nhân.', 'Mẫu 2020-2022 nhỏ hơn nhiều so với 2018 (974, 961 vs 3.721).', 'K10 và MHC-SF — không tách lo âu riêng.']:
        add_p(doc, f'• {b}')

    add_page_ref(doc, '2–4', 'EpiPsychSci', 'Vol. 34, 2025')
    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Tỷ lệ căng thẳng tâm lý cao ở giới trẻ Úc ngày càng tăng. Khảo sát Quốc gia về SKTT và Hạnh phúc (NSMHW) 2020-2022 cho thấy 39,6% trẻ 16–24 tuổi đã trải nghiệm rối loạn tâm thần trong 12 tháng qua. Tuy nhiên, NC hiện tại chủ yếu tập trung vào bệnh tật, ít chú ý đến SKTT tích cực.')
    add_p(doc, 'Mô hình Sức khỏe Tâm thần Toàn diện (CMH) của Keyes (2002, 2005) nhận ra bệnh tâm thần và hạnh phúc là HAI CHIỀU KHÁC BIỆT — không phải hai đầu của một thang. Một người có thể có căng thẳng nhưng vẫn hạnh phúc (struggling), hoặc không bệnh nhưng không hạnh phúc (languishing).')
    add_p(doc, '6 trạng thái CMH: Flourishing (hạnh phúc cao + không bệnh), Middling (hạnh phúc trung bình + không bệnh), Languishing (hạnh phúc thấp + không bệnh), Struggling (hạnh phúc cao + có bệnh), Stumbling (hạnh phúc trung bình + có bệnh), Floundering (hạnh phúc thấp + có bệnh).')

    add_page_ref(doc, '3–4', 'EpiPsychSci', 'Vol. 34, 2025')
    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Thiết kế cắt ngang lặp lại. 3 khảo sát quốc gia: 2018 (n=3.721), 2020 (n=974), 2022 (n=961). Phỏng vấn điện thoại (CATI) qua Kantar Public. Mẫu quota: tuổi, giới, tiểu bang đại diện dân số Úc. Chỉ phân tích dữ liệu CATI (loại online) để nhất quán phương pháp.')
    add_p(doc, 'Thang đo:', bold=True)
    add_p(doc, '• K10 (Kessler-10): 10 mục, đo căng thẳng tâm lý. Phân loại: Thấp (10–15), Trung bình (16–21), Cao (22–29), Rất cao (30–50).')
    add_p(doc, '• MHC-SF (Mental Health Continuum Short Form): 14 mục, đo hạnh phúc chủ quan, tâm lý, xã hội. Phân loại: Thấp, Trung bình, Cao (flourishing).')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Trạng thái CMH theo năm khảo sát', 3)
    add_table(doc,
        ['Trạng thái', '2018 (%)', '2020 (%)', '2022 (%)', 'Xu hướng'],
        [['Flourishing', '53,0', '52,2', '44,4', '↓ GIẢM'], ['Middling', '16,0', '14,7', '14,6', '— Ổn định'], ['Languishing', '0,3', '0,4', '0,3', '— Rất thấp'], ['Struggling', '9,1', '11,3', '10,2', '— Ổn định'], ['Stumbling', '19,3', '19,0', '27,3', '↑ TĂNG MẠNH'], ['Floundering', '2,4', '2,5', '3,2', '↑ Tăng nhẹ']],
        widths=[3.0, 2.0, 2.0, 2.0, 3.0])
    add_p(doc, 'χ²(10) = 44,59; p < 0,001. Flourishing 2022 thấp hơn 2018 (z = 4,76, p < 0,001) và 2020 (z = 3,41, p = 0,001).', size=9, italic=True)

    add_heading(doc, 'Bảng 2. Flourishing theo giới tính và năm', 3)
    add_table(doc,
        ['Năm', 'Nam Flourishing', 'Nữ Flourishing', 'Chênh lệch', 'p'],
        [['2018', '58,3%', '47,9%', '10,4 điểm', '<0,001'], ['2020', '57,6%', '46,7%', '10,9 điểm', '0,001'], ['2022', '52,2%', '36,3%', '15,9 điểm', '0,001']],
        widths=[2.0, 3.0, 3.0, 3.0, 2.0])
    add_p(doc, 'Khoảng cách giới MỞ RỘNG: từ 10,4 điểm (2018) lên 15,9 điểm (2022).', size=9, italic=True)

    add_heading(doc, 'Bảng 3. Flourishing theo nhóm tuổi (2018, n=3.721)', 3)
    add_table(doc,
        ['Nhóm tuổi', 'Flourishing %', 'Stumbling %', 'So với 12–14'],
        [['12–14', '70,7', '9,4', 'Reference'], ['15–17', '51,9', '20,3', 'p < 0,001'], ['18–21', '43,3', '24,4', 'p < 0,001'], ['22–25', '47,9', '22,2', 'p < 0,001']],
        widths=[3.0, 3.0, 3.0, 3.0])
    add_p(doc, 'Flourishing GIẢM mạnh từ 12–14 tuổi (70,7%) đến 18–21 tuổi (43,3%), sau đó tăng nhẹ ở 22–25 (47,9%).', size=9, italic=True)

    add_p(doc, 'K10 cao/rất cao:', bold=True)
    add_p(doc, '• 2018: 30,7% (cao 20,7% + rất cao 10,0%)')
    add_p(doc, '• 2020: 32,7% (24,0% + 8,7%)')
    add_p(doc, '• 2022: 40,7% (27,2% + 13,5%) — TĂNG RÕ RỆT')

    add_heading(doc, '4. THẢO LUẬN', 2)
    add_p(doc, 'SKTT giới trẻ Úc đang xấu đi: flourishing giảm, stumbling tăng, K10 cao/rất cao tăng từ 30,7% lên 40,7%. Khoảng cách giới mở rộng — nữ bị ảnh hưởng nặng hơn. Mô hình CMH cho thấy bức tranh phức tạp hơn chỉ đo bệnh: nhiều người "struggling" (có bệnh nhưng vẫn hạnh phúc) — gợi ý hạnh phúc có thể là yếu tố bảo vệ.')

    add_heading(doc, '5. KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu 3 khảo sát quốc gia Úc cho thấy SKTT toàn diện của giới trẻ xấu đi đáng kể từ 2018 đến 2022, với flourishing giảm từ 53% xuống 44,4% và căng thẳng cao/rất cao tăng lên 40,7%, gợi ý rằng cần tiếp cận CMH toàn diện (đo cả bệnh tật và hạnh phúc) thay vì chỉ tập trung vào pathology.')

    add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
    add_p(doc, 'Crisp, D. et al. (2025). The complete mental health of Australia\'s adolescents and emerging adults. EpiPsychSci, 34, e16.', size=10)
    add_p(doc, 'Keyes, C.L.M. (2002). The mental health continuum. Journal of Health and Social Behavior, 43, 207–222.', size=10)
    add_p(doc, '(Xem đầy đủ trong bài gốc — 40+ TLTK)', size=10, italic=True)

    add_abbreviation_table(doc, [('K10','Kessler-10 — Thang Căng thẳng Tâm lý 10 mục'),('MHC-SF','Mental Health Continuum Short Form — Thang Liên tục SKTT Rút gọn'),('CMH','Complete Mental Health — Sức khỏe Tâm thần Toàn diện'),('CATI','Computer Assisted Telephone Interview — Phỏng vấn Điện thoại Hỗ trợ Máy tính'),('VTN','Vị thành niên'),('SKTT','Sức khỏe tâm thần')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['EpiPsychSci Q1 IF≈7. Khung CMH toàn diện — đo cả bệnh VÀ hạnh phúc.','3 khảo sát quốc gia (2018-2022), phương pháp nhất quán.','Phân tích theo tuổi, giới, trạng thái CMH — rất chi tiết.','Phát hiện khoảng cách giới mở rộng — nữ bị ảnh hưởng nặng hơn theo thời gian.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế chi tiết:', bold=True)
    for s in ['Chỉ Úc — quốc gia phương Tây phát triển, khác VN.','Mẫu 2020 và 2022 nhỏ hơn nhiều (974, 961 vs 3.721) — sức mạnh thống kê giảm.','K10 đo căng thẳng chung — không tách lo âu riêng.','CATI (phỏng vấn điện thoại) — có thể thiên lệch so với online.','Gender diverse quá ít (<2%) để phân tích.','Cắt ngang lặp lại — không theo dõi cùng cá nhân.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
    for s in ['VN CHƯA CÓ khảo sát CMH toàn diện (đo cả bệnh + hạnh phúc).','Cần áp dụng khung CMH cho dữ liệu VN — VTN VN flourishing hay stumbling?','So sánh xu hướng Úc vs VN: flourishing giảm ở cả hai?']:
        add_red(doc, f'• {s}')

    save(doc, '25_EpiPsychSci_2025.docx')

# =====================================================================
# QT26 — UK NHS Parliament 2024
# =====================================================================
def dich_QT26():
    print('QT26 UK NHS...')
    doc = create_doc()
    add_p(doc, 'Link: https://commonslibrary.parliament.uk/research-briefings/sn06988/', size=10)
    add_heading(doc, 'Thống kê sức khỏe tâm thần: Tỷ lệ, dịch vụ và tài trợ tại Anh', 1)
    h = doc.add_paragraph(); r = h.add_run('Mental Health Statistics: Prevalence, Services and Funding in England'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Mental Health Statistics: Prevalence, Services and Funding in England'),
        ('Tác giả', 'Carl Baker, Esme Kirk-Wade'),
        ('Cơ quan', 'UK Parliament — House of Commons Library'),
        ('Xuất bản', '1 March 2024, Briefing Paper SN06988, 46 trang'),
        ('Loại tài liệu', 'Báo cáo thống kê chính sách (briefing paper) — KHÔNG phải nghiên cứu gốc'),
        ('Phạm vi', 'England (Anh quốc), dữ liệu NHS'),
    ])
    add_page_ref(doc, '1–46', 'UK Parliament Briefing', 'SN06988, 2024')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Báo cáo tổng hợp thống kê SKTT tại Anh: tỷ lệ rối loạn, dịch vụ NHS, chi tiêu. Dữ liệu chính: 1/6 người trưởng thành có vấn đề SKTT phổ biến. Ở trẻ em (7–16 tuổi): tỷ lệ rối loạn tâm thần có thể xác định tăng từ 11,6% (2017) lên 17,8% (2021) rồi giảm nhẹ xuống 16,3% (2023). Lượt giới thiệu SKTT trẻ em/VTN đến NHS tăng mạnh.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['UK Parliament — nguồn dữ liệu chính sách uy tín.','Tỷ lệ rối loạn SKTT trẻ em 7–16 tăng: 11,6% (2017) → 17,8% (2021) → 16,3% (2023).','Lo âu riêng ở trẻ 7–16: từ 3,5% (2017) lên 6,3% (2023) — vẫn cao hơn trước COVID.','Giới thiệu SKTT VTN đến NHS: 1,07 triệu lượt (2023-24).','Chi tiêu SKTT NHS: £16,3 tỷ/năm — 10,6% tổng ngân sách NHS.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Chỉ England — không đại diện toàn UK.','Briefing paper — tổng hợp, không phải NC gốc.','Bối cảnh NHS (phổ quát, miễn phí) khác VN hoàn toàn.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. TỶ LỆ RỐI LOẠN SKTT TRẺ EM', 2)
    add_heading(doc, 'Bảng 1. Tỷ lệ rối loạn SKTT ở trẻ em Anh (NHS Digital surveys)', 3)
    add_table(doc,
        ['Chỉ số', '2017', '2020', '2021', '2022', '2023', 'Xu hướng'],
        [['Rối loạn SKTT (7–16)', '11,6%', '—', '17,8%', '18,0%', '16,3%', 'Tăng rồi giảm nhẹ'],
         ['Rối loạn SKTT (17–19)', '10,1%', '—', '17,4%', '—', '—', 'Tăng mạnh'],
         ['Lo âu (7–16)', '3,5%', '—', '—', '—', '6,3%', '↑ Tăng gần gấp đôi'],
         ['Rối loạn cảm xúc (7–16)', '4,1%', '—', '—', '—', '7,8%', '↑ Tăng']],
        widths=[3.5, 1.5, 1.5, 1.5, 1.5, 1.5, 3.0])

    add_heading(doc, 'Bảng 2. Dịch vụ SKTT trẻ em/VTN NHS', 3)
    add_table(doc,
        ['Chỉ số', 'Giá trị', 'Ghi chú'],
        [['Giới thiệu đến CYPMHS', '1,07 triệu (2023-24)', 'Tăng đều mỗi năm'],
         ['Thời gian chờ trung bình', 'Nhiều tháng', 'Khác nhau theo vùng'],
         ['Chi tiêu SKTT NHS', '£16,3 tỷ/năm', '10,6% tổng ngân sách NHS'],
         ['Tỷ lệ tiếp cận dịch vụ', 'Cao (NHS phổ quát)', 'So VN: chỉ 8,4%']],
        widths=[4.0, 4.0, 5.0])

    add_heading(doc, '2. XU HƯỚNG VÀ PHÂN TÍCH', 2)
    add_p(doc, 'Xu hướng tỷ lệ SKTT trẻ em Anh song song với xu hướng toàn cầu: tăng trước COVID, tăng mạnh trong COVID, giảm nhẹ sau COVID nhưng vẫn cao hơn trước đại dịch. Lo âu ở trẻ 7–16 tăng gần gấp đôi (3,5% → 6,3%).')
    add_p(doc, 'Chênh lệch giới tính: nữ bị ảnh hưởng nặng hơn nam, đặc biệt ở nhóm 17–19 tuổi.')
    add_p(doc, 'Dịch vụ: NHS phổ quát nhưng thời gian chờ dài, nhu cầu vượt cung. 1,07 triệu lượt giới thiệu/năm cho CYPMHS (Children and Young People\'s Mental Health Services).')

    add_heading(doc, '3. KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu NHS England cho thấy tỷ lệ rối loạn SKTT trẻ em tăng đáng kể từ 2017 đến 2023, với lo âu tăng gần gấp đôi. Mặc dù NHS phổ quát, nhu cầu vẫn vượt cung. So sánh với VN (chỉ 8,4% tiếp cận) cho thấy khoảng cách khổng lồ trong dịch vụ SKTT VTN.')

    add_abbreviation_table(doc, [('NHS','National Health Service — Dịch vụ Y tế Quốc gia Anh'),('CYPMHS','Children and Young People\'s Mental Health Services'),('SKTT','Sức khỏe tâm thần'),('VTN','Vị thành niên')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['UK Parliament — nguồn chính thức, dữ liệu NHS đáng tin cậy.','Xu hướng dài hạn với nhiều thời điểm (2017-2023).','Bao gồm cả tỷ lệ, dịch vụ và chi tiêu — bức tranh toàn diện.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['Briefing paper — tổng hợp, không phải NC gốc. Không có peer review.','Chỉ England, không đại diện toàn UK (Scotland, Wales, NI khác).','NHS phổ quát — bối cảnh hoàn toàn khác VN (y tế tư nhân, bảo hiểm hạn chế).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN cần báo cáo thống kê tương tự — tổng hợp tỷ lệ + dịch vụ + chi tiêu SKTT VTN.','So sánh chi tiêu SKTT: UK £16,3 tỷ vs VN chưa có con số.']:
        add_red(doc, f'• {s}')

    save(doc, '26_UK_NHS_2025_Parliament.docx')

# Run
dich_QT25()
dich_QT26()
print('Done QT25-26!')
