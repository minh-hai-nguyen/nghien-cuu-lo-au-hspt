# -*- coding: utf-8 -*-
"""
Build VN002 VNAMHS 2022 FULL translation — Jenkins template, Option C coverage 0.90+.

Structure:
1. Cover + title bilingual
2. Bibliographic info table
3. Abstract (TÓM TẮT) translated
4. Quick review (TÓM TẮT ĐÁNH GIÁ NHANH)
5. Abbreviations table
6. Table of contents
7. Full body translation (51 pages)
8. References (English, verbatim)
9. 7-section critique (red)

Target:
- Coverage ratio 0.85-0.95
- 12/12 images embedded with bilingual captions
- Page markers --- Trang X ---
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_imgs/'

doc = Document()

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(11)

# Helpers -----------------------------------------------------
def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False, underline=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if underline: r.underline = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def H1(text):
    return P(text, bold=True, size=16, align=WD_ALIGN_PARAGRAPH.CENTER,
             color=RGBColor(0x1F, 0x3A, 0x68))

def H2(text):
    return P(text, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))

def H3(text):
    return P(text, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))

def H4(text):
    return P(text, bold=True, italic=True, size=11)

def PageMarker(page_num, note='V-NAMHS, UNICEF/IOS 2022'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'--- Trang {page_num}, {note} ---')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

def AddImg(filename, caption_vn, caption_en, width_cm=11.0):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path):
        print(f'  [SKIP] Image not found: {filename}')
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(path, width=Cm(width_cm))
    except Exception as e:
        print(f'  [ERR] {filename}: {e}')
        return
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2', first_col_width=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(10)
        r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
    return t

# =============================================================
# 1. COVER
# =============================================================
H1('KHẢO SÁT SỨC KHOẺ TÂM THẦN VỊ THÀNH NIÊN VIỆT NAM')
H1('(V-NAMHS)')
P('BÁO CÁO VỀ CÁC PHÁT HIỆN CHÍNH', bold=True, size=14, align=WD_ALIGN_PARAGRAPH.CENTER)
P('Viet Nam Adolescent Mental Health Survey — Report on Main Findings',
  italic=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
P()
P('THÁNG 11 NĂM 2022 | NOVEMBER 2022', bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
P()
P('Canonical ID: VN002 | Mã dự án Lo âu: VN002_VNAMHS_2022_National',
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER,
  color=RGBColor(0x80, 0x80, 0x80))
P()

# Cover image (V-NAMHS logo/hero)
AddImg('p001_img4_Image27.jpg',
       'Ảnh bìa — Báo cáo V-NAMHS 2022',
       'Cover image — V-NAMHS 2022 Report',
       width_cm=12.0)

P()
P('Trích dẫn đề xuất (Suggested Citation): Institute of Sociology, University of Queensland, and Johns Hopkins Bloomberg School of Public Health. 2022. Viet Nam Adolescent Mental Health Survey: Report on Main Findings. Hanoi, Viet Nam: Institute of Sociology.',
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
P()

# =============================================================
# 2. BIBLIOGRAPHIC INFO TABLE
# =============================================================
H2('THÔNG TIN THƯ MỤC')
info_rows = [
    ('Tiêu đề tiếng Việt', 'Khảo sát Sức khoẻ Tâm thần Vị thành niên Việt Nam — Báo cáo về các Phát hiện chính'),
    ('Tiêu đề tiếng Anh', 'Viet Nam Adolescent Mental Health Survey (V-NAMHS) — Report on Main Findings'),
    ('Năm công bố', '2022 (tháng 11)'),
    ('Tác giả chính', 'Vũ Mạnh Lợi, Nguyễn Đức Vinh (Viện Xã hội học); Đào Thị Khánh Hoa (Hội Xã hội học VN); Holly E. Erskine, Cartiah McGrath, Krystina Wallis, Sarah J. Blondell, Harvey A. Whiteford, James G. Scott (Đại học Queensland); Robert Blum, Shoshanna Fine, Mengmeng Li, Astha Ramaiya (Johns Hopkins BSPH)'),
    ('Cơ quan xuất bản', 'Viện Xã hội học (IOS) thuộc Viện Hàn lâm KHXH Việt Nam (VASS), phối hợp với The University of Queensland (UQ) và Johns Hopkins Bloomberg School of Public Health (JHSPH)'),
    ('Loại tài liệu', 'Báo cáo điều tra dân số — quốc gia (nationally representative household survey)'),
    ('Ngôn ngữ gốc', 'Tiếng Anh'),
    ('Số trang bản gốc', '51 trang (~132.000 ký tự EN)'),
    ('Địa bàn khảo sát', '38 tỉnh/thành phố thuộc 4 vùng địa lý, bao phủ cả đô thị và nông thôn; đại diện quốc gia VTN 10–17 tuổi'),
    ('Cỡ mẫu', '5.996 cặp phụ huynh–vị thành niên (81,1 % tỷ lệ trả lời từ 6.048 hộ đủ tiêu chuẩn)'),
    ('Thời gian thực địa', '21/9/2021 đến 16/12/2021, bắt đầu miền Bắc xuống miền Nam'),
    ('Công cụ đo chính', 'DISC-5 (Diagnostic Interview Schedule for Children, Version 5) — chẩn đoán theo DSM-5'),
    ('Rối loạn đo lường', '6 loại: ám ảnh xã hội, rối loạn lo âu lan toả, rối loạn trầm cảm nặng, PTSD, rối loạn hành vi, ADHD'),
    ('Khung hợp tác quốc tế', 'NAMHS — cùng Indonesia (I-NAMHS) và Kenya (K-NAMHS)'),
    ('Ngày dịch/cập nhật', '14/04/2026 — rebuild Jenkins template từ bản v1 (06/04/2026)'),
    ('Canonical ID', 'VN002 (dự án Lo âu)'),
]
tbl = doc.add_table(rows=len(info_rows), cols=2)
tbl.style = 'Table Grid'
for i, (k, v) in enumerate(info_rows):
    c1 = tbl.rows[i].cells[0]
    c2 = tbl.rows[i].cells[1]
    c1.text = ''; c2.text = ''
    r1 = c1.paragraphs[0].add_run(k)
    r1.font.name = 'Times New Roman'; r1.font.size = Pt(10); r1.bold = True
    set_cell_bg(c1, 'E7E6E6')
    r2 = c2.paragraphs[0].add_run(v)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(10)
doc.add_paragraph()

# =============================================================
# 3. ABSTRACT — TÓM TẮT
# =============================================================
H2('TÓM TẮT (Abstract)')

P('Bối cảnh. ', bold=True)
P('Hiểu biết về tỷ lệ mắc rối loạn sức khoẻ tâm thần ở vị thành niên Việt Nam còn rất hạn chế. Các nghiên cứu sẵn có chủ yếu dùng thang sàng lọc (SDQ, CBCL, YSR) trên mẫu nhỏ hoặc hẹp địa bàn, khiến việc khái quát lên quy mô quốc gia và so sánh quốc tế rất khó khăn. V-NAMHS là khảo sát đầu tiên ở Việt Nam sử dụng công cụ chẩn đoán tiêu chuẩn quốc tế (DISC-5) để ước tính tỷ lệ rối loạn tâm thần ở vị thành niên 10–17 tuổi theo DSM-5 trên mẫu đại diện quốc gia.')

P('Phương pháp. ', bold=True)
P('Điều tra cắt ngang hộ gia đình tại 38 tỉnh/thành phố thuộc 4 vùng (Trung du miền núi phía Bắc + Tây Nguyên; Đồng bằng sông Hồng; Bắc Trung Bộ + Duyên hải miền Trung; Đông Nam Bộ + Đồng bằng sông Cửu Long). Chọn ngẫu nhiên 200 địa bàn điều tra (EA) phân đều 50 EA/vùng (25 đô thị + 25 nông thôn). Tổng 7.599 hộ được tiếp cận, 6.048 hộ đủ tiêu chuẩn và đồng ý tham gia (tỷ lệ trả lời 81,1 %), loại trừ 52 hộ có dữ liệu không đầy đủ; mẫu cuối cùng 5.996 cặp phụ huynh–vị thành niên. Phụ huynh được đo PSC-17, PHQ-9, GAD-7 và trả lời về ADHD của con. Vị thành niên được phỏng vấn bằng 5 module DISC-5 (ám ảnh xã hội, lo âu lan toả, trầm cảm nặng, rối loạn hành vi, PTSD), cùng bộ câu hỏi về hành vi tự sát, tự làm hại bản thân, bắt nạt, gia đình, peer, học đường, COVID-19. Tỷ lệ và số lượng được gia trọng (weighted) theo phân bố tuổi–giới–đô thị/nông thôn của quần thể VTN 10–17 tuổi Việt Nam.')

P('Kết quả chính. ', bold=True)
P('(1) 21,7 % VTN có ít nhất một vấn đề sức khoẻ tâm thần (mental health problem) trong 12 tháng qua (n = 1.301), trong đó lo âu phổ biến nhất (18,6 %), tiếp theo trầm cảm (4,3 %), ADHD (2,8 %), PTSD (1,0 %), rối loạn hành vi (0,7 %). (2) 3,3 % đáp ứng đầy đủ tiêu chí chẩn đoán DSM-5 cho ít nhất một rối loạn tâm thần trong 12 tháng qua (n = 200); rối loạn lo âu là phổ biến nhất (2,3 %), tiếp theo trầm cảm nặng (0,9 %), ADHD (0,5 %), PTSD (0,3 %), rối loạn hành vi (0,2 %). Không có khác biệt tổng thể theo giới tính hay nhóm tuổi 10–13 vs 14–17. (3) Lĩnh vực suy giảm chức năng (impairment) do triệu chứng tập trung cao nhất ở gia đình (67,0 % trong nhóm có vấn đề SKTT; 74,2 % trong nhóm có rối loạn), tiếp theo bạn bè (47,0 %; 63,3 %) và trường/việc làm (45,4 %; 64,1 %). (4) Trong 12 tháng qua, 1,4 % VTN có ý nghĩ tự sát, 0,4 % lập kế hoạch tự sát, 0,2 % toan tự sát; 1,6 % từng toan tự sát trong đời; 1 % tự làm hại bản thân trong 12 tháng và 4,7 % từng tự làm hại trong đời. Trên 70 % các em có hành vi tự sát 12 tháng qua đồng thời có vấn đề SKTT. (5) Khoảng trống điều trị lớn: chỉ 8,4 % các em có vấn đề SKTT (n = 109) và 6,5 % toàn mẫu đã sử dụng bất kỳ dịch vụ nào trong 12 tháng qua. Bác sĩ/điều dưỡng là nơi đến nhiều nhất (56,2 %). Chỉ 5,1 % phụ huynh nhận ra con mình cần trợ giúp; trong nhóm không tìm kiếm dịch vụ, rào cản hàng đầu là "tự giải quyết trong gia đình" (20,4 %), "không chắc con có cần" (10,7 %) và "không biết chỗ tìm trợ giúp" (10,0 %). 73,9 % VTN tâm sự với người thân trong gia đình khi có lo lắng; 38,2 % với bạn bè. (6) Tác động COVID-19: 7,7 % VTN thường xuyên trải qua ≥ 1 vấn đề cảm xúc/hành vi nhiều hơn bình thường trong đại dịch (nếu gộp cả "đồng ý" lên tới 67 %). 71,5 % hộ gia đình báo cáo giảm thu nhập trong đại dịch; 12,3 % thường xuyên thiếu tiền cho nhu yếu phẩm. Trong 7,1 % phụ huynh báo cáo con cần trợ giúp cảm xúc/hành vi trong đại dịch, 80,3 % không tiếp cận được dịch vụ, chủ yếu do sợ lây nhiễm COVID-19 (69,2 %).')

P('Hàm ý. ', bold=True)
P('V-NAMHS cung cấp lần đầu dữ liệu đại diện quốc gia về tỷ lệ rối loạn SKTT ở VTN Việt Nam theo chuẩn DSM-5, khẳng định SKTT là vấn đề y tế công cộng quan trọng. Cần chính sách quốc gia toàn diện về SKTT VTN, ưu tiên: lồng ghép sàng lọc SKTT vào chăm sóc ban đầu, nâng cao mental health literacy cho phụ huynh, đào tạo giáo viên/y tế trường học về sàng lọc và giới thiệu chuyển tuyến, bảo đảm dịch vụ dễ tiếp cận (hotline, tư vấn trực tuyến) đặc biệt trong khủng hoảng như đại dịch.')

P()
P('(Tóm tắt do người dịch soạn dựa trên Executive Summary + toàn văn báo cáo; mọi con số đã verify vs PDF gốc — tham khảo PHẢN BIỆN Vòng 2 cuối bài.)',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))
doc.add_paragraph()

# =============================================================
# 4. QUICK REVIEW
# =============================================================
H2('TÓM TẮT ĐÁNH GIÁ NHANH')
P('Điểm mạnh:', bold=True)
P('• Khảo sát đại diện quốc gia ĐẦU TIÊN của Việt Nam về rối loạn SKTT ở VTN theo chẩn đoán DSM-5/DISC-5')
P('• Cỡ mẫu rất lớn (n = 5.996) trên 38 tỉnh bao phủ 4 vùng địa lý, urban + rural')
P('• Phỏng vấn trực tiếp face-to-face bởi 127 phỏng vấn viên được đào tạo, có monitoring UQ+JHSPH')
P('• Tỷ lệ trả lời cao (81,1 %) dù diễn ra trong đại dịch COVID-19')
P('• Công cụ được adapt văn hoá (translation/back-translation + review lâm sàng) cho bối cảnh Việt Nam')
P('• Uy tín thể chế cao: VASS + UQ + JHSPH; khung NAMHS ba nước (Indonesia + Kenya + VN)')
P('• Đồng thời báo cáo 2 mức — "mental health problem" (half-threshold) và "mental disorder" (full DSM-5) — giúp nhận diện nhóm subthreshold quan trọng cho can thiệp sớm')

P('Hạn chế:', bold=True)
P('• Thiết kế cắt ngang, không xác lập nhân quả')
P('• Phỏng vấn trực tiếp có thể underestimate hành vi tự sát và các vấn đề stigmatised (so với self-administered questionnaire)')
P('• DISC-5 dựa trên DSM-5 chủ yếu là Western criteria; có thể không bắt hết somatic expression của VTN Việt Nam (Kim et al. 2019)')
P('• Câu hỏi dịch vụ chủ yếu hỏi phụ huynh, thiếu góc nhìn VTN về barriers to care')
P('• 18–19 tuổi bị loại (DISC-5 không phù hợp), nên kết quả không áp dụng cho nhóm late adolescence')
P('• Data thu thập giữa COVID-19 có thể tạo "period effect"; vẫn cần so sánh với wave hậu đại dịch')

P('Hướng cải thiện:', bold=True)
P('• Lặp lại theo chu kỳ 3–5 năm để theo dõi xu hướng và đánh giá hiệu quả chính sách')
P('• Bổ sung module somatic complaints culturally adapted')
P('• Phân tích subgroup sâu hơn: dân tộc thiểu số, LGBTQ, VTN khuyết tật')
P('• Nghiên cứu longitudinal nhỏ hơn để truy nguồn nhân quả risk factors')
doc.add_paragraph()

# =============================================================
# 5. ABBREVIATIONS
# =============================================================
H2('BẢNG VIẾT TẮT')
abbr_rows = [
    ('ADHD', 'Attention-deficit/hyperactivity disorder — Rối loạn tăng động giảm chú ý'),
    ('CBCL', 'Child Behaviour Checklist — Bảng kiểm Hành vi Trẻ em'),
    ('DISC-5', 'Diagnostic Interview Schedule for Children, Version 5 — Phỏng vấn chẩn đoán DSM-5 cho trẻ em, phiên bản 5'),
    ('DSM-5', 'Diagnostic and Statistical Manual of Mental Disorders, 5th Edition — Cẩm nang chẩn đoán và thống kê rối loạn tâm thần, tái bản thứ 5'),
    ('EA', 'Enumeration Area — Địa bàn điều tra'),
    ('GBD', 'Global Burden of Disease Study — Nghiên cứu gánh nặng bệnh tật toàn cầu'),
    ('GEAS', 'Global Early Adolescent Study — Nghiên cứu VTN sớm toàn cầu'),
    ('GOPFP', 'General Office for Population and Family Planning — Tổng cục Dân số Kế hoạch hoá gia đình'),
    ('GSHS', 'Global School-based Student Health Survey — Khảo sát Sức khoẻ Học sinh toàn cầu'),
    ('GSO', 'General Statistical Office — Tổng cục Thống kê Việt Nam'),
    ('HSPI', 'Health Strategy and Policy Institute — Viện Chiến lược & Chính sách Y tế'),
    ('HICs', 'High-income countries — Các nước thu nhập cao'),
    ('HREC', 'Human Research Ethics Committee — Hội đồng Đạo đức Nghiên cứu trên người'),
    ('I-NAMHS', 'Indonesia National Adolescent Mental Health Survey'),
    ('IOS', 'Institute of Sociology — Viện Xã hội học (VASS)'),
    ('JHSPH', 'Johns Hopkins Bloomberg School of Public Health'),
    ('K-NAMHS', 'Kenya National Adolescent Mental Health Survey'),
    ('LMICs', 'Low- and middle-income countries — Các nước thu nhập thấp và trung bình'),
    ('MOH', 'Ministry of Health — Bộ Y tế'),
    ('NAMHS', 'National Adolescent Mental Health Surveys'),
    ('POPFP', 'Provincial Office of Population and Family Planning — Chi cục DS-KHHGĐ tỉnh'),
    ('PTSD', 'Posttraumatic stress disorder — Rối loạn căng thẳng sau sang chấn'),
    ('SDQ', 'Strengths and Difficulties Questionnaire — Bảng hỏi Điểm mạnh và Điểm yếu'),
    ('SESD', 'Social-Environmental Statistics Department — Vụ Thống kê Xã hội và Môi trường'),
    ('UQ', 'The University of Queensland — Đại học Queensland (Úc)'),
    ('VASS', 'Viet Nam Academy of Social Sciences — Viện Hàn lâm KHXH Việt Nam'),
    ('V-NAMHS', 'Viet Nam Adolescent Mental Health Survey — Khảo sát SKTT VTN Việt Nam'),
    ('WHO', 'World Health Organization — Tổ chức Y tế Thế giới'),
    ('YSR', 'Youth Self-Report — Tự báo cáo Vị thành niên'),
    ('MOLISA', 'Ministry of Labour, Invalids and Social Affairs — Bộ Lao động Thương binh Xã hội'),
    ('MOET', 'Ministry of Education and Training — Bộ Giáo dục và Đào tạo'),
]
MakeTable(['Viết tắt', 'Giải nghĩa (EN — VN)'], abbr_rows, header_bg='D9E1F2')
doc.add_paragraph()

# =============================================================
# 6. TABLE OF CONTENTS
# =============================================================
H2('MỤC LỤC')
toc = [
    'GIỚI THIỆU (trang 9)',
    '  • Bối cảnh (Background) — trang 9',
    '  • V-NAMHS là gì? (What is V-NAMHS?) — trang 10',
    '  • Ai tổ chức V-NAMHS? (Who conducted?) — trang 10',
    '  • Ai tham gia V-NAMHS? (Who participated?) — trang 10',
    '  • Người tham gia được hỏi những gì? — trang 12',
    '  • Khi nào V-NAMHS được tiến hành? — trang 12',
    '  • COVID-19 đã tác động như thế nào? — trang 12',
    '  • Phạm vi của báo cáo — trang 12',
    'ĐẶC ĐIỂM MẪU (SAMPLE CHARACTERISTICS) — trang 14',
    '  • Vị thành niên — trang 14',
    '  • Phụ huynh — trang 15',
    'SỨC KHOẺ TÂM THẦN (MENTAL HEALTH) — trang 17',
    '  • Tổng quan — trang 17',
    '  • Đo lường — trang 18',
    '  • Phát hiện — trang 20',
    '  • Suy xét (Interpretation, Limitations, Implications) — trang 23',
    'SỬ DỤNG DỊCH VỤ (SERVICE USE) — trang 25',
    '  • Tổng quan — trang 25',
    '  • Đo lường — trang 26',
    '  • Phát hiện — trang 26',
    '  • Suy xét — trang 29',
    'COVID-19 — trang 32',
    '  • Tổng quan — trang 32',
    '  • Đo lường — trang 33',
    '  • Phát hiện — trang 33',
    '  • Suy xét — trang 34',
    'PHỤ LỤC 1: Các công cụ đo lường (Measures) — trang 36',
    'PHỤ LỤC 2: Phương pháp (Methodology) — trang 38',
    'PHỤ LỤC 3: Thuật ngữ (Glossary) — trang 41',
    'PHỤ LỤC 4: Nhóm nghiên cứu (Research teams) — trang 44',
    'TÀI LIỆU THAM KHẢO (References) — trang 45 (giữ tiếng Anh, APA)',
    'PHẢN BIỆN VÀ ĐÁNH GIÁ (Người dịch) — 7 phần chi tiết + meta-review 2 vòng',
]
for line in toc:
    p = doc.add_paragraph()
    r = p.add_run(line)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)

doc.add_paragraph()
print(f'Part 1 (front matter) written. Paragraphs so far: ~150')

# Save checkpoint
doc.save(OUT)
print(f'Saved checkpoint to {OUT}')
