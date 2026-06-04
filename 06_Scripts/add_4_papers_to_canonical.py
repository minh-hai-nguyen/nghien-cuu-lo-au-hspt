"""Build 4 Tom-tat docx + add vao canonical_index.json + update RAG/KG.

4 bai:
- QT068: Urao 2018 BMC — Journey of the Brave 10 sessions × 45 min Japan
- QT069: Urao 2022 BMC — Journey of the Brave 14 sessions × 20 min Japan
- QT070: Moore 2017 World J Psychiatry — Meta bullying victimization OR=1,77 anxiety
- QT071: Matsumoto K 2024 JMIR Pediatrics — Unguided iCBT smartphone SAD adolescents Japan
"""
import sys, io, json, copy, re
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

TOMTAT_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/Tom-tat-tung-bai')
CANONICAL = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/02_Papers-goc/canonical_index.json')
LIGHT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/tro-ly-nghien-cuu-tam-ly-light/web/data')
HEAVY = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/tro-ly-nghien-cuu-tam-ly/web/data')


# Build Tom-tat docx
def build_tomtat(filename, title, paragraphs):
    out = TOMTAT_DIR / filename
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(14)

    for label, text in paragraphs:
        p = d.add_paragraph()
        if label:
            r = p.add_run(label); r.bold = True; r.font.size = Pt(12); r.add_break()
        r = p.add_run(text); r.font.size = Pt(12)

    d.save(out)
    print(f'  ✓ {out.name}')
    return out


# 4 papers data
PAPERS = {
    'QT068': {
        'descriptor': 'Urao_2018_JapanCBT_Journey_BMC',
        'filename': 'QT068_Urao_2018_JapanCBT_BMC.docx',
        'title': 'QT068 — Tóm tắt: Urao và cộng sự (2018) Journey of the Brave Japan',
        'tomtat': [
            ('Tên công trình:', '"Effectiveness of a cognitive behavioural therapy-based anxiety prevention programme at an elementary school in Japan: a quasi-experimental study" — Urao, Yoshida, Sato, Shimizu (2018), BMC Child and Adolescent Psychiatry and Mental Health (DOI: 10.1186/s13034-018-0240-5).'),
            ('Phương pháp:', 'Quasi-experimental tại trường tiểu học Nhật. Mẫu: n=72 (41 intervention + 31 control), tuổi 10-11 (lớp 5). Can thiệp "Journey of the Brave" 10 sessions × 45 phút, 2 buổi/tháng × 6 tháng (10/2013–3/2014). Nhóm chứng: chương trình thông thường.'),
            ('Công cụ:', 'Spence Children\'s Anxiety Scale (SCAS) — primary; Strengths and Difficulties Questionnaire (SDQ) — secondary.'),
            ('Nội dung CBT:', 'Tâm lý giáo dục lo âu + kỹ năng thư giãn + tái cấu trúc nhận thức + kỹ năng quyết đoán.'),
            ('Kết quả:', 'SCAS post-intervention: chênh nhóm −5,321 (KTC 95% −10,12 đến −0,523; p=0,030). 3-month follow-up: SCAS −7,104 (p=0,004); SDQ −3,284 (p=0,002). Hiệu quả DUY TRÌ sau can thiệp.'),
            ('Đánh giá chất lượng:', '⭐⭐⭐ Trung bình–Khá. Quasi-experimental (không hoàn toàn ngẫu nhiên), cỡ mẫu nhỏ (n=72), nhưng hiệu quả ý nghĩa thống kê và duy trì 3 tháng. BMC Q1.'),
            ('Ý nghĩa cho VN:', 'Mô hình CBT 10 sessions phù hợp lứa tuổi cuối cấp 1 / đầu THCS. Có thể adapt cho HS THCS Việt Nam (lớp 6).'),
        ],
        'topics': ['CBT school-based', 'Anxiety prevention', 'Japan adolescents'],
        'region': 'Khác',
        'method': 'Quasi-experimental',
    },
    'QT069': {
        'descriptor': 'Urao_2022_JourneyBrave_14sessions_BMC',
        'filename': 'QT069_Urao_2022_JourneyBrave_14sessions_BMC.docx',
        'title': 'QT069 — Tóm tắt: Urao và cộng sự (2022) Journey of the Brave 14 sessions',
        'tomtat': [
            ('Tên công trình:', '"School-based cognitive behavioural intervention programme for addressing anxiety in 10- to 11-year-olds using short classroom activities in Japan: a quasi-experimental study" — Urao, Yoshida, Sato, Shimizu (2022), BMC Psychiatry (DOI: 10.1186/s12888-022-04326-y).'),
            ('Phương pháp:', 'Quasi-experimental Nhật. Mẫu: n=90 (31 intervention + 59 control), tuổi 10-11. Can thiệp 14 weekly sessions × 20 phút (TỔNG 5 GIỜ ngắn hơn QT068). Format: short classroom activities.'),
            ('Công cụ:', 'SCAS (primary) + SDQ (secondary).'),
            ('Kết quả:', '2-month follow-up: SCAS Cohen d=0,46 (medium effect; chênh nhóm −5,77, p=0,009); SDQ d=0,15 (small; chênh −1,60, p=0,024).'),
            ('So sánh QT068 vs QT069:', 'QT068 (2018): 10 sessions × 45 phút (7,5h tổng); SCAS chênh −5,32. QT069 (2022): 14 sessions × 20 phút (5h tổng); SCAS d=0,46. → Format NGẮN-CỰC NGẮN có thể HIỆU QUẢ tương đương — phù hợp tích hợp vào lịch sinh hoạt lớp.'),
            ('Đánh giá chất lượng:', '⭐⭐⭐⭐ Cao. Quasi-experimental nhưng cỡ mẫu vừa, hiệu quả medium effect size, BMC Psychiatry Q1. Format ngắn-cực ngắn phù hợp thực tế trường học.'),
            ('Ý nghĩa cho VN:', 'Mô hình 14 sessions × 20 phút (tổng 5h) RẤT PHÙ HỢP cho sinh hoạt lớp/đầu giờ tại trường THCS Việt Nam — không chiếm nhiều thời gian học chính khoá.'),
        ],
        'topics': ['CBT school-based', 'Brief intervention', 'Japan adolescents'],
        'region': 'Khác',
        'method': 'Quasi-experimental',
    },
    'QT070': {
        'descriptor': 'Moore_2017_BullyingMeta_WorldJPsy',
        'filename': 'QT070_Moore_2017_BullyingMeta_WorldJPsy.docx',
        'title': 'QT070 — Tóm tắt: Moore và cộng sự (2017) Meta bullying → mental health',
        'tomtat': [
            ('Tên công trình:', '"Consequences of bullying victimization in childhood and adolescence: A systematic review and meta-analysis" — Moore SE, Norman RE, Suetani S, Thomas HJ, Sly PD, Scott JG (2017), World Journal of Psychiatry, 7(1), 60–76 (DOI: 10.5498/wjp.v7.i1.60).'),
            ('Phương pháp:', 'Systematic review và meta-analysis với quality-effects model. 165 bài (từ 317 đánh giá ban đầu): 57 prospective cohort + 108 cross-sectional. Đánh giá nhân quả theo Bradford Hill criteria và World Cancer Research Fund grading.'),
            ('Kết quả CỐT LÕI cho lo âu:', 'Bullying victimization → Anxiety: OR = 1,77 (KTC 95% 1,34–2,33). Đây là META-ANALYSIS NỀN TẢNG được trích dẫn rộng rãi cho liên kết bullying → anxiety.'),
            ('Kết quả khác:', 'Depression OR=2,21 (1,34–3,65); Suicidal ideation OR=1,77 (1,56–2,02); Self-injury OR=1,75 (1,40–2,19); Tobacco OR=1,36 (0,96–1,92, NS).'),
            ('Đánh giá nhân quả:', 'Bằng chứng MẠNH NHẤT cho quan hệ nhân quả: depression, anxiety, poor general health, suicidal ideation. Phân tích tuổi: <13 vs >13 cho thấy bullying có hiệu ứng tương tự ở cả hai nhóm.'),
            ('Đánh giá chất lượng:', '⭐⭐⭐⭐⭐ Rất cao. Meta-analysis 165 bài với phân tích chất lượng nghiêm ngặt; được trích dẫn hơn 1.000 lần; nguồn số liệu chuẩn cho liên kết bullying → mental health.'),
            ('Ý nghĩa cho VN:', 'Tham chiếu CHÍNH cho phản biện về phát hiện β BNHĐ → RLLACL = 0,376 trong chương 3. Cụ thể: OR Moore 2017 = 1,77 cho anxiety nói chung; chương 3 phát hiện β=0,376 đặc thù cho lo âu chia ly — bổ sung cho y văn quốc tế.'),
        ],
        'topics': ['Meta-analysis', 'Bullying victimization', 'Adolescent mental health'],
        'region': 'Khác',
        'method': 'Meta-analysis',
    },
    'QT071': {
        'descriptor': 'MatsumotoK_2024_iCBT_SAD_JMIR',
        'filename': 'QT071_MatsumotoK_2024_iCBT_SAD_JMIR.docx',
        'title': 'QT071 — Tóm tắt: Matsumoto K (2024) Unguided iCBT smartphone SAD',
        'tomtat': [
            ('Tên công trình:', '"Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: Multicenter Randomized Controlled Trial" — Matsumoto K, Hamatani S, Shiga K, và cộng sự (2024), JMIR Pediatrics and Parenting (DOI: 10.2196/55786).'),
            ('Phương pháp:', 'Multicenter RCT Nhật. Mẫu: n=77 (31 intervention + 38 control phân tích). Tuổi 15-25. Can thiệp: 10 sessions text-based qua smartphone app (learningBOX), HOÀN TOÀN UNGUIDED. Nội dung dựa trên Clark & Wells cognitive model: psychoeducation, safety behaviors, video feedback, behavioral experiments, relapse prevention.'),
            ('Công cụ:', 'Liebowitz Social Anxiety Scale (LSAS) primary. Treatment response: ≥28% LSAS giảm. Remission: LSAS <35.'),
            ('Kết quả CHÍNH:', 'Response rate: 61% intervention vs 24% control (OR=4,97; p=0,003). Recovery: 68% vs 34% (OR=3,95; p=0,008). LSAS Hedge g=0,64–0,66 (medium effect). Dropout chỉ 9%.'),
            ('Phát hiện ĐẶC BIỆT:', 'TRÁI với Walder 2025 phân tích tổng hợp 21 RCT DMHI cho rằng UNGUIDED có hiệu lực YẾU (g≈0,2): bài này UNGUIDED nhưng đạt g=0,64–0,66 (medium). Có thể do (1) format SMARTPHONE app dễ tiếp cận hơn web; (2) thiết kế SAD-specific theo Clark & Wells; (3) tuổi 15-25 phù hợp với smartphone usage.'),
            ('Đánh giá chất lượng:', '⭐⭐⭐⭐ Cao. RCT đa trung tâm, JMIR Q1, dropout rate thấp 9%, hiệu quả medium. Cỡ mẫu nhỏ (n=77) là hạn chế.'),
            ('Ý nghĩa cho VN:', 'Mô hình thiết kế tiềm năng cho VN: smartphone app tiếng Việt cho SAD học sinh THCS-THPT. SAD-specific + Clark & Wells model + 10 sessions text-based. KHÔNG cần hướng dẫn người (giảm chi phí). Phù hợp đề xuất Hành động 4 trong giả thuyết NĐT.'),
        ],
        'topics': ['iCBT', 'Smartphone app', 'Social anxiety disorder', 'Japan adolescents'],
        'region': 'Khác',
        'method': 'RCT',
    },
}


def main():
    # Step 1: Build Tom-tat
    print('=== Build Tom-tat ===')
    for qt_id, data in PAPERS.items():
        build_tomtat(data['filename'], data['title'], data['tomtat'])

    # Step 2: Add to canonical
    print('\n=== Update canonical_index ===')
    with open(CANONICAL, encoding='utf-8') as f:
        can = json.load(f)

    added_can = 0
    for qt_id, data in PAPERS.items():
        if qt_id not in can:
            can[qt_id] = {
                'descriptor': data['descriptor'],
                'summary': data['filename'],
                'pdf_folder': 'The-gioi_Au-My-Uc' if 'Japan' not in data['descriptor'] else 'Khac',
            }
            added_can += 1

    with open(CANONICAL, 'w', encoding='utf-8') as f:
        json.dump(can, f, ensure_ascii=False, indent=2)
    qt_count = len([k for k in can if k.startswith('QT')])
    vn_count = len([k for k in can if k.startswith('VN')])
    print(f'  +{added_can} entries; total {qt_count} QT + {vn_count} VN = {qt_count+vn_count}')

    # Step 3: Update KG + RAG
    print('\n=== Update KG + RAG ===')
    new_papers_kg = [
        {'id': 'PAPER_QT068_Urao_2018', 'type': 'Paper',
         'label': 'QT068 Urao (2018) BMC Journey of the Brave Japan 10 sessions'},
        {'id': 'PAPER_QT069_Urao_2022', 'type': 'Paper',
         'label': 'QT069 Urao (2022) BMC Psychiatry 14 sessions × 20 min'},
        {'id': 'PAPER_QT070_Moore_2017_BullyingMeta', 'type': 'Paper',
         'label': 'QT070 Moore (2017) Meta-analysis bullying → anxiety OR=1,77'},
        {'id': 'PAPER_QT071_MatsumotoK_2024_iCBT', 'type': 'Paper',
         'label': 'QT071 Matsumoto K (2024) JMIR Unguided iCBT smartphone SAD'},
    ]

    new_concepts = [
        {'id': 'CONCEPT_JOURNEY_OF_THE_BRAVE', 'type': 'Concept', 'label': 'Journey of the Brave (CBT Nhật adolescents)'},
        {'id': 'CONCEPT_BRIEF_CBT_FORMAT', 'type': 'Concept', 'label': 'Brief CBT format (14 sessions × 20 min ngắn-cực ngắn)'},
        {'id': 'CONCEPT_BULLYING_ANXIETY_OR', 'type': 'Concept', 'label': 'Bullying → anxiety OR (Moore 2017: 1,77)'},
        {'id': 'CONCEPT_SCAS', 'type': 'Concept', 'label': 'SCAS (Spence Children Anxiety Scale)'},
        {'id': 'CONCEPT_LSAS', 'type': 'Concept', 'label': 'LSAS (Liebowitz Social Anxiety Scale)'},
        {'id': 'CONCEPT_CLARK_WELLS_MODEL', 'type': 'Concept', 'label': 'Clark & Wells cognitive model SAD'},
    ]

    new_edges = [
        # Paper - Concept
        ('PAPER_QT068_Urao_2018', 'CONCEPT_JOURNEY_OF_THE_BRAVE', 'RELATED_TO'),
        ('PAPER_QT068_Urao_2018', 'CONCEPT_SCAS', 'RELATED_TO'),
        ('PAPER_QT069_Urao_2022', 'CONCEPT_JOURNEY_OF_THE_BRAVE', 'RELATED_TO'),
        ('PAPER_QT069_Urao_2022', 'CONCEPT_BRIEF_CBT_FORMAT', 'RELATED_TO'),
        ('PAPER_QT070_Moore_2017_BullyingMeta', 'CONCEPT_BULLYING_ANXIETY_OR', 'RELATED_TO'),
        ('PAPER_QT070_Moore_2017_BullyingMeta', 'CONCEPT_SCHOOL_BULLYING', 'RELATED_TO'),
        ('PAPER_QT071_MatsumotoK_2024_iCBT', 'CONCEPT_iCBT', 'RELATED_TO'),
        ('PAPER_QT071_MatsumotoK_2024_iCBT', 'CONCEPT_LSAS', 'RELATED_TO'),
        ('PAPER_QT071_MatsumotoK_2024_iCBT', 'CONCEPT_CLARK_WELLS_MODEL', 'RELATED_TO'),
        ('PAPER_QT071_MatsumotoK_2024_iCBT', 'CONCEPT_GUIDED_VS_UNGUIDED_DMHI', 'RELATED_TO'),
    ]

    for label, dir_ in [('LIGHT', LIGHT), ('HEAVY', HEAVY)]:
        print(f'\n  ==== {label} ====')
        # KG
        path = dir_ / 'questions_kg.json'
        with open(path, encoding='utf-8') as f: kg = json.load(f)
        eids = {n['id'] for n in kg['nodes']}
        eedges = {(e['source'],e['target'],e['relation']) for e in kg['edges']}
        na = 0
        for n in new_papers_kg + new_concepts:
            if n['id'] not in eids: kg['nodes'].append(n); eids.add(n['id']); na += 1
        ea = 0
        for src,tgt,rel in new_edges:
            if (src,tgt,rel) in eedges or src not in eids or tgt not in eids: continue
            kg['edges'].append({'source':src,'target':tgt,'relation':rel}); eedges.add((src,tgt,rel)); ea += 1
        from collections import Counter
        kg['meta']['n_nodes']=len(kg['nodes']); kg['meta']['n_edges']=len(kg['edges'])
        kg['meta']['node_types']=dict(Counter(n['type'] for n in kg['nodes']))
        kg['meta']['edge_relations']=dict(Counter(e['relation'] for e in kg['edges']))
        kg['meta']['updated']='2026-05-09'
        with open(path,'w',encoding='utf-8') as f: json.dump(kg,f,ensure_ascii=False,indent=2)
        nn = len(kg['nodes']); ne = len(kg['edges'])
        print(f'    KG: +{na} nodes, +{ea} edges; total {nn}/{ne}')

    print('\nDone.')


if __name__ == '__main__':
    main()
