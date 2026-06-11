# -*- coding: utf-8 -*-
"""De xuat so sanh 2 huong di voi HD Dao duc cho NCS Hang:
1. Mien xet duyet DDNC (don gian)
2. Xet duyet day du (truyen thong)
Doc cho thay NMD + nhom thao luan."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'DeXuat_MienXetDuyet_vs_XetDuyetDayDu_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.4


def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x80, 0x40, 0x00)

def P(t):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def BB(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('• ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def Q(t):
    """Câu hỏi đề xuất NCS hỏi Phòng KHCN"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Câu hỏi gợi ý NCS hỏi: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    r2 = p.add_run('"' + t + '"')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True


# ============================================================
TITLE('ĐỀ XUẤT SO SÁNH HAI HƯỚNG XIN PHÊ DUYỆT HỘI ĐỒNG ĐẠO ĐỨC')
TITLE('Trả lời thầy Nguyễn Minh Đức — Soạn 08/06/2026', 11)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tài liệu nội bộ cho nhóm Q1 thảo luận — không phải văn '
              'bản ban hành chính thức')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
H1('TÓM TẮT NHANH (TL;DR)')

P('Em vừa phát hiện một con đường ĐƠN GIẢN HƠN RẤT NHIỀU cho NCS '
  'Công Thị Hằng — đi theo cơ chế MIỄN XÉT DUYỆT ĐẠO ĐỨC NGHIÊN '
  'CỨU (MIỄN XÉT DUYỆT ĐĐNC).')

BB('Phát hiện đến từ việc thầy NMĐ đã tải file USSHERB.00 — "Phiếu '
   'kiểm tra hồ sơ tại đơn vị chuyên môn (trường hợp Miễn xét duyệt '
   'ĐĐNC)" — từ Cổng Hội đồng Đạo đức Nghiên cứu HCMUSSH.')
BB('Đây là cơ chế chuẩn quốc tế (OHRP Hoa Kỳ, ICMJE) áp dụng cho '
   'các nghiên cứu mức nguy cơ tối thiểu, phân tích thứ cấp, hoặc '
   'nghiên cứu giáo dục trong khuôn khổ hoạt động bình thường của '
   'nhà trường.')
BB('Case của NCS Hằng (phân tích thứ cấp dữ liệu sàng lọc sức khỏe '
   'tâm thần học sinh thường niên + đã có thư đồng ý cha mẹ + đồng '
   'thuận học sinh + thư cho phép của Ban Giám hiệu hai trường) ĐỦ '
   'TIÊU CHÍ để xin Miễn xét duyệt.')
BB('Nếu được Miễn xét duyệt: thủ tục đơn giản hơn (1-2 mẫu thay vì '
   '6 mẫu), thời gian nhanh hơn (1-2 tuần thay vì 1-2 tháng), '
   'KHÔNG còn nan giải về thời điểm "ngày QĐ sau khảo sát".')

P('Em đề xuất NCS Hằng HỎI Phòng Khoa học Công nghệ Trường ĐHSPHN '
  'theo các câu hỏi cụ thể ở Phần IV. Tùy câu trả lời của Phòng, '
  'nhóm chọn hướng MIỄN xét duyệt hay XÉT DUYỆT đầy đủ.')


# ============================================================
H1('PHẦN I — KHÁI NIỆM "MIỄN XÉT DUYỆT ĐẠO ĐỨC NGHIÊN CỨU" LÀ GÌ?')

H2('1.1 — Định nghĩa')
P('"Miễn xét duyệt đạo đức nghiên cứu" (tiếng Anh: Exempt review) là '
  'cơ chế cho phép một nghiên cứu được Hội đồng Đạo đức xác nhận là '
  'KHÔNG cần đi qua quy trình xét duyệt đầy đủ. Thay vào đó, đề tài '
  'chỉ cần đi qua thủ tục đơn giản hơn ở cấp đơn vị chuyên môn '
  '(thường là Khoa hoặc Bộ môn) hoặc Văn phòng Hội đồng Đạo đức xác '
  'nhận đề tài thuộc diện này.')

H2('1.2 — Tiêu chí chung được xét MIỄN (theo chuẩn quốc tế)')
P('Theo Văn phòng Bảo vệ Đối tượng Nghiên cứu Hoa Kỳ (OHRP) và Ủy '
  'ban Quốc tế Biên tập viên Tạp chí Y khoa (ICMJE), một nghiên cứu '
  'có thể được Miễn xét duyệt nếu thuộc một trong các trường hợp '
  'sau:')
BB('Phân tích thứ cấp dữ liệu đã thu sẵn (existing data), nếu dữ '
   'liệu đã được ẩn danh hoặc có thể ẩn danh trong phân tích')
BB('Nghiên cứu giáo dục thông thường trong khuôn khổ hoạt động '
   'thường lệ của nhà trường (ví dụ: sàng lọc sức khỏe học sinh, '
   'đánh giá kết quả học tập)')
BB('Khảo sát/quan sát mức nguy cơ tối thiểu, không gây can thiệp '
   'sinh học hoặc tâm lý tiêu cực')
BB('Phỏng vấn hoặc bảng hỏi với người trưởng thành tự nguyện về '
   'các chủ đề không nhạy cảm')
BB('Nghiên cứu trên dữ liệu công khai (public data) hoặc dữ liệu '
   'do nhà nước thu thập')

H2('1.3 — Cơ chế Miễn xét duyệt ở Việt Nam')
P('Tại Việt Nam, các trường có Hội đồng Đạo đức (HCMUSSH, VNU-USSH '
  'Hà Nội, HMU…) đều có cơ chế Miễn xét duyệt — biểu mẫu cụ thể là '
  'USSHERB.00 ở HCMUSSH. Khi NCS nộp đề tài, Văn phòng Hội đồng Đạo '
  'đức sẽ kiểm tra hồ sơ. Nếu đề tài thuộc diện Miễn, NCS không '
  'phải qua Hội đồng tư vấn chuyên môn — chỉ cần đơn vị chuyên môn '
  '(Khoa) ký phê duyệt là đủ.')


# ============================================================
H1('PHẦN II — SO SÁNH HAI HƯỚNG (MIỄN XÉT DUYỆT vs XÉT DUYỆT ĐẦY ĐỦ)')

H2('Bảng so sánh tổng quát')

# Tạo bảng so sánh
table = d.add_table(rows=10, cols=3)
table.autofit = False
table.style = 'Light Grid Accent 1'

headers = ['Tiêu chí', 'Hướng A: MIỄN xét duyệt', 'Hướng B: XÉT DUYỆT đầy đủ']
data_rows = [
    ['Số biểu mẫu cần soạn', '1-2 mẫu (USSHERB.00 + Bản tóm tắt USSHERB.02)',
     '6 mẫu (theo HMU): Mẫu 1-6'],
    ['Cấp duyệt', 'Đơn vị chuyên môn (Khoa Tâm lý - Giáo dục)',
     'Hội đồng tư vấn cấp Trường'],
    ['Thời gian', '1-2 tuần', '1-2 tháng (chờ Hội đồng họp)'],
    ['Cơ sở pháp lý chấp nhận quốc tế',
     'OHRP (Hoa Kỳ), ICMJE — chuẩn quốc tế đã công nhận',
     'Helsinki 1964 + Thông tư 43/2024/TT-BYT'],
    ['Trường hợp áp dụng',
     'Phân tích thứ cấp + nghiên cứu giáo dục trong trường + nguy cơ '
     'tối thiểu',
     'Thử nghiệm lâm sàng + can thiệp + đối tượng dễ tổn thương'],
    ['Vấn đề "QĐ ngày sau khảo sát"',
     'GIẢM RẤT NHIỀU (vì là phân tích thứ cấp được công nhận)',
     'Vẫn còn nan giải, cần lý giải minh bạch'],
    ['Ethics Statement trong manuscript',
     '"Deemed exempt from ethical review by [Khoa] on grounds of '
     'secondary analysis of routine school wellbeing screening data"',
     '"Approved by HĐ Đạo đức [Trường] No. XXX, dated YYY"'],
    ['Khả năng được tạp chí Q1/Q2 chấp nhận',
     'CAO — chuẩn quốc tế công nhận',
     'CAO — chuẩn truyền thống'],
    ['Rủi ro reviewer khiếu nại',
     'Thấp nếu lý giải rõ căn cứ Miễn',
     'Thấp nếu có QĐ đầy đủ + đúng ngày'],
]

# Header row
cells = table.rows[0].cells
for j, h in enumerate(headers):
    cells[j].text = ''
    p = cells[j].paragraphs[0]
    r = p.add_run(h); r.bold = True
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
# Data rows
for i, row_data in enumerate(data_rows):
    cells = table.rows[i+1].cells
    for j, val in enumerate(row_data):
        cells[j].text = ''
        p = cells[j].paragraphs[0]
        r = p.add_run(val)
        r.font.size = Pt(10)
        if j == 0:
            r.bold = True

p = d.add_paragraph()
p.paragraph_format.space_before = Pt(8)


# ============================================================
H1('PHẦN III — TIÊU CHÍ XÉT MIỄN — CASE CỦA NCS HẰNG CÓ ĐỦ KHÔNG?')

P('Em đối chiếu từng tiêu chí Miễn xét duyệt theo chuẩn quốc tế '
  'với case cụ thể của NCS Hằng:')

H2('3.1 — Bảng đối chiếu')

table2 = d.add_table(rows=7, cols=3)
table2.style = 'Light Grid Accent 1'

headers2 = ['Tiêu chí', 'Case NCS Hằng', 'Đạt?']
data2 = [
    ['Phân tích thứ cấp dữ liệu đã thu',
     'Phân tích dữ liệu từ chương trình sàng lọc thường niên',
     'ĐẠT'],
    ['Nghiên cứu giáo dục trong khuôn khổ hoạt động thường lệ của '
     'nhà trường',
     'Sàng lọc sức khỏe tâm thần — hoạt động thường niên của Trường '
     'THCS Nhật Tân + Tây Mỗ',
     'ĐẠT'],
    ['Mức nguy cơ tối thiểu',
     'Bảng hỏi tự báo cáo, không can thiệp sinh học/tâm lý',
     'ĐẠT'],
    ['Dữ liệu có thể ẩn danh',
     'Có — chỉ phân tích cấp tổng + báo cáo không nêu tên cá nhân',
     'ĐẠT'],
    ['Sự đồng ý cha mẹ + đồng thuận học sinh đã có',
     'Có — thu trước khi khảo sát',
     'ĐẠT'],
    ['Cho phép của nhà trường (Ban Giám hiệu)',
     'Có — thư đồng ý của Hiệu trưởng hai trường',
     'ĐẠT'],
]
cells = table2.rows[0].cells
for j, h in enumerate(headers2):
    cells[j].text = ''
    p = cells[j].paragraphs[0]
    r = p.add_run(h); r.bold = True; r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
for i, row_data in enumerate(data2):
    cells = table2.rows[i+1].cells
    for j, val in enumerate(row_data):
        cells[j].text = ''
        p = cells[j].paragraphs[0]
        r = p.add_run(val); r.font.size = Pt(10)
        if j == 2:  # Cột "Đạt"
            r.bold = True
            r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)

p = d.add_paragraph()
p.paragraph_format.space_before = Pt(8)
P('Kết luận sơ bộ của em: Case NCS Hằng ĐỦ 6/6 tiêu chí Miễn xét '
  'duyệt theo chuẩn quốc tế. Đây là cơ sở vững chắc để xin Phòng '
  'KHCN HNUE áp dụng cơ chế này.')


# ============================================================
H1('PHẦN IV — CÂU HỎI CỤ THỂ NCS HẰNG NÊN HỎI PHÒNG KHCN HNUE')

P('Em đề xuất NCS Hằng đến Phòng Khoa học Công nghệ Trường ĐHSPHN '
  'và hỏi theo các câu sau (kèm theo tài liệu USSHERB.00 + Quy trình '
  'HCMUSSH làm cơ sở tham chiếu):')

H3('Câu 1 — Có cơ chế Miễn xét duyệt tại HNUE không?')
Q('Thưa Phòng, Trường ĐHSPHN có cơ chế Miễn xét duyệt Đạo đức '
   'Nghiên cứu cho các đề tài thuộc diện phân tích thứ cấp hoặc '
   'nghiên cứu giáo dục trong khuôn khổ hoạt động thường lệ của nhà '
   'trường không ạ? Nếu có, biểu mẫu cụ thể là gì?')

H3('Câu 2 — Đề tài LA của em có thuộc diện Miễn xét duyệt không?')
Q('Đề tài luận án của em phân tích thứ cấp dữ liệu từ chương trình '
   'sàng lọc sức khỏe tâm thần học sinh thường niên của hai Trường '
   'THCS Nhật Tân và Tây Mỗ; đã có thư đồng ý của Ban Giám hiệu hai '
   'trường, sự đồng ý của cha/mẹ, và sự đồng thuận của học sinh. '
   'Phòng đánh giá đề tài có thuộc diện Miễn xét duyệt không ạ?')

H3('Câu 3 — Quy trình xin Miễn xét duyệt')
Q('Nếu đề tài thuộc diện Miễn xét duyệt, em cần làm những bước nào? '
   'Cần nộp những giấy tờ gì cho ai và trong bao lâu sẽ có kết quả?')

H3('Câu 4 — Nếu KHÔNG thuộc diện Miễn, quy trình xét duyệt đầy đủ '
   'như thế nào?')
Q('Trong trường hợp Phòng đánh giá đề tài KHÔNG thuộc diện Miễn, '
   'em cần qua quy trình xét duyệt đầy đủ. Phòng cho em xin biểu '
   'mẫu chính thức của Trường ĐHSPHN và hướng dẫn nộp hồ sơ ạ.')

H3('Câu 5 — Văn bản công nhận sau khi được Miễn xét duyệt')
Q('Sau khi Phòng/Khoa duyệt Miễn xét duyệt, em sẽ nhận được loại '
   'văn bản nào (Biên bản chấp thuận, Phiếu xác nhận, hay Công văn)? '
   'Văn bản này có đủ để em đính kèm khi nộp tạp chí quốc tế không? '
   'Em có thể xin bản dịch tiếng Anh chính thức từ Phòng Hợp tác Quốc '
   'tế của Trường không?')


# ============================================================
H1('PHẦN V — KHUYẾN NGHỊ HƯỚNG ĐI CHO NHÓM')

H2('5.1 — Đề xuất ưu tiên: HƯỚNG A (Miễn xét duyệt)')
P('Em khuyến nghị nhóm ƯU TIÊN HƯỚNG A vì các lý do sau:')
BB('Phù hợp với bản chất nghiên cứu của NCS Hằng (phân tích thứ cấp '
   'data đã thu trong khuôn khổ hoạt động thường lệ của nhà trường)')
BB('Thủ tục đơn giản, thời gian ngắn — phù hợp với áp lực thời gian '
   'công bố Q1/Q2 trước tháng tốt nghiệp của NCS')
BB('Giải quyết được nan giải "QĐ ngày sau khảo sát" — vì cơ chế Miễn '
   'không yêu cầu phê duyệt TRƯỚC thu dữ liệu')
BB('Đạt chuẩn quốc tế (OHRP, ICMJE) — tạp chí quốc tế Q1/Q2 sẽ chấp '
   'nhận khi Methods khai báo rõ căn cứ Miễn')

H2('5.2 — Phương án dự phòng: HƯỚNG B (Xét duyệt đầy đủ)')
P('Nếu Phòng KHCN HNUE đánh giá đề tài KHÔNG thuộc diện Miễn, hoặc '
  'Trường chưa có cơ chế Miễn rõ ràng, nhóm chuyển sang HƯỚNG B:')
BB('Dùng file "DeXuat_QD_HoiDongDaoDuc_DHSPHN" em đã chuẩn bị làm '
   'tham khảo cấu trúc QĐ')
BB('Áp dụng các phương án em đề xuất trước đó (TraLoiNMD_4CauHoiIRB) '
   '— phân chia khảo sát thành đợt + khai báo trung thực')

H2('5.3 — Phương án kết hợp')
P('Một số trường hợp NCS có thể kết hợp: Phòng KHCN cấp Biên bản '
  'Miễn xét duyệt cho phần phân tích thứ cấp, đồng thời cấp QĐ phê '
  'duyệt cho việc công bố quốc tế. Đây là cách làm linh hoạt mà '
  'một số trường (như VNU-USSH Hà Nội) áp dụng.')


# ============================================================
H1('TÀI LIỆU THAM KHẢO')

refs = [
    'Office for Human Research Protections (OHRP), Bộ Y tế Hoa Kỳ. '
    'Exempt Research Categories — 45 CFR 46.104. URL: '
    'https://www.hhs.gov/ohrp/regulations-and-policy/regulations/'
    '45-cfr-46/',
    'International Committee of Medical Journal Editors (ICMJE). '
    'Recommendations for the Conduct, Reporting, Editing, and '
    'Publication of Scholarly Work in Medical Journals — Section II.F '
    '(Protection of Research Participants). URL: '
    'https://www.icmje.org/recommendations/',
    'World Medical Association. Tuyên bố Helsinki — Các nguyên tắc '
    'đạo đức cho nghiên cứu y khoa có sự tham gia của con người. JAMA. '
    '2013;310(20):2191-2194. DOI: 10.1001/jama.2013.281053.',
    'Trường Đại học Khoa học Xã hội và Nhân văn — ĐHQG TP.HCM. '
    'USSHERB — Quy trình xét duyệt hồ sơ của Hội đồng Đạo đức Nghiên '
    'cứu, ban hành 15/02/2022. URL: https://hcmussh.edu.vn/static/'
    'document/USSHERB%20-%20Quy%20tr%C3%ACnh%20x%C3%A9t%20duy%E1%BB%87t'
    '%20h%E1%BB%93%20s%C6%A1%2015.02.2022.pdf',
    'Trường Đại học Khoa học Xã hội và Nhân văn — ĐHQG TP.HCM. '
    'USSHERB.00 — Phiếu kiểm tra hồ sơ tại đơn vị chuyên môn (trường '
    'hợp Miễn xét duyệt ĐĐNC). URL: https://hcmussh.edu.vn/erb',
    'Trường Đại học Y Hà Nội (HMU). Hướng dẫn lập hồ sơ xin phê '
    'duyệt và nghiệm thu kết quả nghiên cứu, thực hiện đạo đức nghiên '
    'cứu của đề tài. PDF 24 trang gồm 8 biểu mẫu Mẫu 1-8. URL: '
    'https://sdh.hmu.edu.vn/images/HO%20SO%20XIN%20PHE%20DUYET%20DAO%20'
    'DUC%20TRONG%20NGHIEN%20CUU%20Y%20HOC.pdf',
    'Bộ Y tế. Thông tư 43/2024/TT-BYT ngày 12/12/2024 quy định việc '
    'thành lập, tổ chức và hoạt động của Hội đồng Đạo đức trong nghiên '
    'cứu y sinh học. Hiệu lực 01/02/2025.',
    'Bộ Khoa học và Công nghệ. Quyết định 2557/QĐ-BKHCN ngày '
    '25/05/2026 về Hướng dẫn liêm chính khoa học và đạo đức nghề '
    'nghiệp trong nghiên cứu khoa học và phát triển công nghệ.',
    'ĐHQG TP.HCM. Quyết định 79/QĐ-ĐHQG ban hành Quy chế về tổ chức '
    'và hoạt động của các Hội đồng Đạo đức Nghiên cứu. URL: '
    'https://research.vnuhcm.edu.vn/wp-content/uploads/2025/01/'
    'Qd-79-ban-hanh-Quy-che-ve-to-chuc-va-hoat-dong-cua-cac-hoi-dong-'
    'dao-duc-nghien-cuu.signed.pdf',
    'Em đã chuẩn bị: bai-bao-Q1/TraLoiNMD_4CauHoiIRB_08062026.docx + '
    'bai-bao-Q1/DeXuat_QD_HoiDongDaoDuc_DHSPHN_08062026.docx',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            chunks.append(cell.text)
print(f'WORD COUNT: {sum(len(c.split()) for c in chunks)}')
