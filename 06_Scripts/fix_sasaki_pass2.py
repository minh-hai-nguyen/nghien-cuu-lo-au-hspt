# -*- coding: utf-8 -*-
"""Second pass — fix remaining Sasaki references that escaped first pass."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

REPLACEMENTS_PASS2 = [
    ('QT45 Sasaki:', 'QT45 Matsumoto:'),
    ('QT45 Sasaki OR', 'QT45 Matsumoto OR'),
    ('Cross-ref QT045 Sasaki', 'Cross-ref QT045 Matsumoto'),
    ('Mâu thuẫn với Sasaki', 'Mâu thuẫn với Matsumoto'),
    ('của Sasaki, N.', 'của Matsumoto, K.'),
    ('Sasaki 2024', 'Matsumoto 2024'),
    ('Sasaki', 'Matsumoto'),  # catch-all last resort
]

def fix_docx(path, replacements):
    d = Document(path)
    changes = 0
    for p in d.paragraphs:
        t = p.text
        new_t = t
        for find, rep in replacements:
            if find in new_t:
                new_t = new_t.replace(find, rep)
        if new_t != t:
            for run in p.runs: run.text = ''
            if p.runs: p.runs[0].text = new_t
            changes += 1
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    new_t = t
                    for find, rep in replacements:
                        if find in new_t:
                            new_t = new_t.replace(find, rep)
                    if new_t != t:
                        for run in p.runs: run.text = ''
                        if p.runs: p.runs[0].text = new_t
                        changes += 1
    if changes > 0:
        d.save(path)
    return changes

FILES = [
    '01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-5.docx',
    'Tom-tat-tung-bai/QT045_Matsumoto_Japan_iCBT_JMIR_2024.docx',
    'Bai_goc_BaoCao_CanThiep_10042026/Bai_57_QT050_Qiaochu_Wang_Mobile_CBT_2025_BAN_DICH_CHI_TIET.docx',
    'Tom-tat-tung-bai/QT050_Qiaochu_MobileCBT_CPP_2025_ABSTRACT.docx',
]

for fn in FILES:
    path = os.path.join(ROOT, fn)
    if os.path.exists(path):
        n = fix_docx(path, REPLACEMENTS_PASS2)
        print(f'{os.path.basename(fn)}: {n} changes pass 2')

# Verify
print('\nVerification — Sasaki count should be 0:')
for fn in FILES:
    path = os.path.join(ROOT, fn)
    if not os.path.exists(path): continue
    d = Document(path)
    txt = '\n'.join(p.text for p in d.paragraphs)
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                txt += ' ' + c.text
    count = txt.count('Sasaki')
    mcount = txt.count('Matsumoto')
    print(f'  {os.path.basename(fn)}: Sasaki={count}, Matsumoto={mcount}')
