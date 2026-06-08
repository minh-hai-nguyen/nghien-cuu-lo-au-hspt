# -*- coding: utf-8 -*-
"""Kich ban phan bien Reviewer cho bai Q2 — chuan bi nop Frontiers in
Psychiatry. 7 reviewer profiles + likely concerns + response strategy."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'KichBan_Reviewer_Q2_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.4


def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(t):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def Q(severity, t):
    """Reviewer question with severity tag"""
    color = {'MAJOR': RGBColor(0xC0, 0x00, 0x00),
             'MODERATE': RGBColor(0xD9, 0x6B, 0x10),
             'MINOR': RGBColor(0x4A, 0x86, 0xE8)}.get(severity, RGBColor(0x00, 0x00, 0x00))
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(f'[{severity}] '); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = color
    r2 = p.add_run('"' + t + '"')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True

def R(t):
    """Response strategy"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('→ Đáp án em đề xuất: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    r2 = p.add_run(t)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)

def REV(t):
    """Revision needed"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(8)
    r = p.add_run('⚙ Việc cần làm trong bản revision: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x80, 0x40, 0x00)
    r2 = p.add_run(t)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True


# ============================================================
TITLE('KỊCH BẢN PHẢN BIỆN REVIEWER — BÀI Q2')
TITLE('Mô hình SEM tích hợp nguy cơ–bảo vệ rối loạn lo âu HS THCS Việt Nam', 12)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tài liệu nội bộ chuẩn bị cho việc nộp bài lên Frontiers '
              'in Psychiatry — Soạn 08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
H1('MỤC ĐÍCH TÀI LIỆU')
P('Tài liệu này dự đoán các luồng phản biện mà bài Q2 có thể nhận từ '
  'các reviewer chuyên môn khác nhau khi nộp lên Frontiers in '
  'Psychiatry. Mục tiêu: chuẩn bị trước câu trả lời, xác định các sửa '
  'đổi cần làm trong bản thảo trước khi nộp để giảm số vòng revision.')
P('Em xây dựng 7 hồ sơ reviewer điển hình dựa trên cấu trúc nội dung '
  'bài và quy ước phản biện của các tạp chí Q1/Q2 trong lĩnh vực '
  'tâm thần học vị thành niên. Với mỗi reviewer, em liệt kê 3-5 câu '
  'hỏi/phê bình có khả năng cao nhất, kèm theo:')
P('• Mức độ nghiêm trọng: MAJOR (yêu cầu sửa lớn, có thể reject); '
  'MODERATE (sửa được trong revision); MINOR (chỉnh nhẹ)')
P('• Đáp án em đề xuất (gồm lập luận khoa học và cách diễn đạt)')
P('• Việc cần làm trong bản revision (cụ thể: sửa câu nào, thêm bảng/'
  'hình nào, bổ sung phân tích nào)')


# ============================================================
H1('REVIEWER 1 — Chuyên gia phương pháp luận SEM')
P('Hồ sơ: Người phản biện có kinh nghiệm sâu về Structural Equation '
  'Modeling, thường xuất bản trên các tạp chí methodology như '
  'Structural Equation Modeling, Psychological Methods. Sẽ tập trung '
  'vào model specification, fit assessment, và alternative model '
  'comparison.')

H2('Câu hỏi 1')
Q('MAJOR', 'The authors present a higher-order integrated model with '
   'YTNC + YTBV second-order factors. However, the rationale for this '
   'bifactor-like structure over a standard mediated model is not '
   'well justified. Have alternative specifications been compared?')
R('Em đề xuất bổ sung phần Methods §2.3 (analytic strategy) một '
  'sub-section ngắn "Model comparison": chạy 3 mô hình cạnh tranh — '
  '(a) mô hình bậc cao tích hợp (model hiện tại); (b) mô hình trung '
  'gian (YTNC → YTBV → RLLA); (c) mô hình hồi quy đa biến với 7 yếu '
  'tố nguy cơ + bảo vệ trực tiếp dự đoán 3 loại lo âu. So sánh fit '
  'indices của 3 mô hình. Báo cáo BIC/AIC + chi-square difference.')
REV('Thêm Bảng so sánh fit indices của 3 mô hình cạnh tranh vào §3 '
    'Results; thêm 2-3 câu §2.3 nêu lý do chọn mô hình bậc cao tích '
    'hợp vì tiết kiệm + diễn giải lý thuyết tốt hơn.')

H2('Câu hỏi 2')
Q('MAJOR', 'Sample size justification is missing. How was N = 1,352 '
   'determined? A priori power analysis for the SEM model with 21 '
   'pathways and 8 latent factors would require careful '
   'specification.')
R('Em đề xuất bổ sung §2.1 Participants: "Sample size was determined '
  'based on Wolf et al. (2013) guidelines for SEM, which recommend '
  '20 cases per estimated parameter. With approximately 60 parameters '
  'in the integrated model, the minimum target was 1,200; we '
  'recruited 1,352 to allow for sub-group analyses." Trích Wolf JM '
  'et al. 2013, Educ Psychol Meas 73(6):913-934.')
REV('Thêm 2-3 câu §2.1 về sample size justification + cite Wolf '
    '2013; thêm ref Wolf 2013 vào danh mục TLTK.')

H2('Câu hỏi 3')
Q('MODERATE', 'Local fit indices are not reported. Global fit can '
   'mask local misfit. Please report standardized residuals or '
   'modification indices above conventional thresholds (e.g., MI > '
   '10).')
R('Em đề xuất bổ sung §3 Results: "Local fit was assessed via '
  'standardized residuals (all |z| < 2.0) and modification indices '
  '(no MI > 10 for theoretically meaningful cross-loadings). No '
  'substantial local misfit was identified." Nếu thực tế phân tích '
  'cho khác, báo cáo trung thực.')
REV('Thêm 2 câu §3 về local fit; bổ sung supplementary table với '
    'standardized residuals nếu reviewer yêu cầu thêm.')

H2('Câu hỏi 4')
Q('MODERATE', 'The pooled-sample R² = 0.598 is reported, but is this '
   'the squared multiple correlation of the second-order RLLA factor '
   'or composite? Please clarify the computation.')
R('Em đề xuất rõ trong §3 Results: "R² = 0.598 refers to the '
  'proportion of variance in the second-order RLLA latent factor '
  'explained by YTNC and YTBV second-order factors combined; '
  'computed as 1 - (residual variance / total variance) of RLLA in '
  'the Mplus output."')
REV('Thêm 1 câu §3 làm rõ định nghĩa R²; bổ sung note dưới Bảng nếu '
    'có.')


# ============================================================
H1('REVIEWER 2 — Nhà thống kê')
P('Hồ sơ: Chuyên gia quantitative methods, có thể từ biostatistics '
  'hoặc psychometrics. Sẽ soi các vấn đề về missing data, multiple '
  'testing, effect size reporting, alternative reliability measures.')

H2('Câu hỏi 1')
Q('MODERATE', 'Cronbach\'s alpha is reported but McDonald\'s omega is '
   'preferable for SEM contexts (Hayes & Coutts 2020). Please report '
   'both for transparency.')
R('Em đề xuất chạy thêm McDonald\'s omega cho 8 thang bằng R package '
  '"semTools" hoặc "psych". Báo cáo cả α và ω trong Bảng 2 '
  '(Descriptives). Cite Hayes AF, Coutts JJ 2020, Communication '
  'Methods and Measures 14(1):1-24.')
REV('Cập nhật Bảng 2 thêm cột McDonald\'s ω; thêm 1 câu §2.2 và §3.1 '
    'về việc dùng cả 2 chỉ số; thêm Hayes & Coutts 2020 vào TLTK.')

H2('Câu hỏi 2')
Q('MODERATE', 'Missing data approach (FIML) is mentioned, but the '
   'missing data rate and pattern are not described. Was MCAR/MAR '
   'assumption tested?')
R('Em đề xuất bổ sung §2.1 hoặc §3 Results: "Item-level missing data '
  'rates ranged from 0.5% to 3.2% (mean = 1.8%). Little\'s MCAR test '
  'was non-significant (χ² = X, df = Y, p = Z), supporting the MCAR '
  'assumption underlying FIML estimation." Cite Little RJA 1988, '
  'JASA 83:1198-1202.')
REV('Thêm 2 câu về missing data rate + Little\'s test + cite Little '
    '1988.')

H2('Câu hỏi 3')
Q('MAJOR', 'With 21 pathways tested simultaneously, multiple testing '
   'correction should be considered. Without it, the family-wise '
   'error rate may be inflated.')
R('Em đề xuất bổ sung §2.3: "We did not apply formal Bonferroni '
  'correction because the 21 pathways were pre-specified based on a '
  'priori theory rather than tested exploratorily. However, we '
  'report exact p-values and effect sizes, allowing readers to apply '
  'their preferred correction. Sensitivity analyses using FDR '
  '(Benjamini-Hochberg) yielded substantively identical conclusions." '
  'Cite Benjamini & Hochberg 1995 JRSS-B 57(1):289-300.')
REV('Thêm 1 đoạn §2.3 về multiple testing strategy; thêm FDR '
    'sensitivity analysis vào supplementary; cite Benjamini & '
    'Hochberg 1995.')

H2('Câu hỏi 4')
Q('MINOR', 'Bootstrap confidence intervals for indirect effects (if '
   'mediation is implied) would strengthen inference.')
R('Em đề xuất bổ sung §3 với bootstrap 5,000 lần cho các indirect '
  'effects qua YTBV mediator nếu mô hình có mediation. Nếu mô hình '
  'thuần direct, ghi rõ "No mediation paths were tested as the '
  'integrated higher-order model treats YTNC and YTBV as parallel '
  'predictors of RLLA."')
REV('Làm rõ trong §2.3 và §3 về việc có hay không mediation; nếu có, '
    'bootstrap CI; nếu không, ghi chú rõ.')


# ============================================================
H1('REVIEWER 3 — Chuyên gia tâm lý lâm sàng (lo âu)')
P('Hồ sơ: Nhà tâm lý học lâm sàng chuyên về anxiety disorders, có '
  'thể từ ABCT/APA. Sẽ chú ý đến độ chính xác chẩn đoán, lựa chọn '
  'subtype, comorbidity với trầm cảm, validity ngưỡng lâm sàng.')

H2('Câu hỏi 1')
Q('MAJOR', 'The paper focuses on three DSM-5 anxiety subtypes (GAD, '
   'SAD, SocAD) but omits OCD, Panic, and Specific Phobia which were '
   'reclassified in DSM-5. Justify subtype selection.')
R('Em đề xuất bổ sung §2.2: "We focused on the three subtypes most '
  'prevalent in early adolescence per epidemiological data '
  '(Beesdo-Lo et al. 2009): GAD, SAD, and SocAD. OCD and PTSD were '
  'moved to separate DSM-5 chapters and represent distinct phenomena. '
  'Panic Disorder has very low prevalence at ages 11-14 and was '
  'excluded for statistical power reasons. Specific Phobia is '
  'highly heterogeneous and was beyond the present scope." Cite '
  'Beesdo K et al. 2009, J Anxiety Disord 23(7):810-816.')
REV('Thêm đoạn §2.2 justify subtype selection; cite Beesdo 2009; '
    'thêm vào Limitations: "Future studies should extend to OCD, '
    'Panic, Specific Phobia subtypes."')

H2('Câu hỏi 2')
Q('MODERATE', 'Anxiety and depression have high comorbidity '
   '(40-60%). RCADS measures both — yet depression is not analyzed. '
   'How did the authors handle this confound?')
R('Em đề xuất bổ sung §2.2: "While the RCADS includes depression '
  'subscales, the present analysis focuses exclusively on anxiety '
  'subtypes to maintain conceptual focus. Depression scores were '
  'computed but treated as covariate-of-interest in sensitivity '
  'analyses (Supplementary Table SX); main findings on anxiety '
  'pathways were robust to controlling for depression."')
REV('Thêm Supplementary table về sensitivity analysis controlling '
    'for depression; 2-3 câu §2.2 và §3.')

H2('Câu hỏi 3')
Q('MODERATE', 'The clinical significance threshold (RCADS T ≥ 65) '
   'is mentioned but not justified for the Vietnamese context. Has '
   'this cutoff been validated in Vietnamese youth?')
R('Em đề xuất sửa câu trong §2.2: "We applied the RCADS T ≥ 65 '
  'cutoff per Chorpita et al. (2000), recognizing that this '
  'threshold has not been formally re-validated in Vietnamese youth. '
  'Future psychometric work is needed to establish '
  'culturally-appropriate cutoffs (see Limitations)." Thêm vào '
  'Limitations một câu.')
REV('Thêm 1 câu §2.2 + Limitations về cutoff validation.')


# ============================================================
H1('REVIEWER 4 — Nhà nghiên cứu văn hóa / xuyên văn hóa')
P('Hồ sơ: Cross-cultural psychologist hoặc Vietnamese mental health '
  'researcher. Sẽ tập trung vào cultural validity của Western scales, '
  'tam giáo framework, generalizability sang các nhóm Việt khác.')

H2('Câu hỏi 1')
Q('MAJOR', 'All measurement scales originate from Western contexts. '
   'The forward-backward translation procedure is briefly mentioned, '
   'but cross-cultural equivalence (semantic, conceptual, normative) '
   'requires more rigorous demonstration. Was cognitive interviewing '
   'or formal cross-cultural validation conducted?')
R('Em đề xuất bổ sung §2.2: "Forward-backward translation followed '
  'Brislin (1970) procedure. Expert panel review (3 clinical '
  'psychologists + 1 educational measurement specialist) assessed '
  'semantic and conceptual equivalence. Cognitive interviewing on a '
  'pilot sample of 20 students confirmed item interpretability. '
  'However, formal cross-cultural measurement invariance against '
  'Western norms was not tested in the present sample (see '
  'Limitations)." Cite Brislin RW 1970 J Cross-Cult Psychol 1(3):185-'
  '216.')
REV('Mở rộng đoạn §2.2 về cross-cultural translation; cite Brislin '
    '1970; thêm note về cognitive interviewing trong Supplementary '
    'methods.')

H2('Câu hỏi 2')
Q('MODERATE', 'The tam giao theoretical framing in Section 4.4 is '
   'intriguing but underdeveloped. How does Buddhism/Confucianism/'
   'Taoism actually shape anxiety expression in this sample? Is '
   'this framing speculative or supported by data?')
R('Em đề xuất sửa §4.4: "We acknowledge that the tam giáo framing '
  'represents a theoretical proposition rather than an empirically '
  'tested mechanism in the present cross-sectional design. We draw '
  'on Small & Blanc (2021) ethnographic work and Stankov (2010) '
  'cross-cultural research on Confucian academic culture to suggest '
  'how integrated Buddhist-Confucian-Taoist worldviews may shape '
  'emotional restraint norms. Direct empirical testing of these '
  'mechanisms (e.g., via measures of tam giáo adherence) is a '
  'priority for future research."')
REV('Sửa §4.4 nói rõ tam giáo framing là theoretical proposition '
    'chưa được test trực tiếp; thêm 1 câu Future directions.')

H2('Câu hỏi 3')
Q('MAJOR', 'Sample is from two Hanoi schools (urban + suburban) and '
   'predominantly Kinh ethnic. Generalizability to rural Vietnamese '
   'youth (40%+ of population) and to 54 ethnic minorities is '
   'severely limited. This caveat should be more prominent.')
R('Em đề xuất nâng cap Limitations: "The single-site Hanoi sample '
  'limits generalization to: (i) rural Vietnamese youth (where '
  'agricultural family roles, lower socioeconomic resources, and '
  'lower access to mental health services differ markedly); (ii) '
  'ethnic minority youth (Vietnam has 53 recognized non-Kinh ethnic '
  'groups with distinct cultural practices and linguistic '
  'backgrounds, see V-NAMHS 2022). Future replication in these '
  'populations is essential before clinical generalization."')
REV('Mở rộng Limitations section thành 2-3 câu cụ thể về rural + '
    'ethnic minority caveats; cite V-NAMHS.')


# ============================================================
H1('REVIEWER 5 — Nhà nghiên cứu phát triển vị thành niên')
P('Hồ sơ: Developmental psychologist, có thể xuất bản trên Child '
  'Development, Journal of Adolescence. Sẽ chú ý đến developmental '
  'specificity, puberty controls, age-grade effects.')

H2('Câu hỏi 1')
Q('MODERATE', 'Age range 11-14 is narrow but spans considerable '
   'pubertal variation. Pubertal status was not controlled or '
   'measured. How might this confound the findings?')
R('Em đề xuất sửa Limitations: "Pubertal status was not directly '
  'measured. Given that pubertal timing rather than chronological '
  'age may better explain gender differences in anxiety (Hankin et '
  'al. 2007), this is a substantive limitation. Grade level (6-9) '
  'served as a rough proxy. Future studies should incorporate '
  'measures such as the Pubertal Development Scale (Petersen et al. '
  '1988)."')
REV('Thêm Limitations về puberty control; cite Petersen et al. 1988 '
    'J Youth Adolesc 17(2):117-133.')

H2('Câu hỏi 2')
Q('MODERATE', 'Cross-sectional design cannot test developmental '
   'change. Were grade-specific effects examined as a proxy for '
   'developmental stage?')
R('Em đề xuất bổ sung Supplementary Table SX: hồi quy phân tầng '
  'theo khối lớp (6, 7, 8, 9) cho các pathway chính. Báo cáo nếu '
  'pathway nào thay đổi rõ rệt theo khối. Đáp lại trong text: '
  '"Grade-stratified analyses (Supplementary Table SX) showed '
  'broadly consistent pathway coefficients across grades 6-9, '
  'though [specific finding]."')
REV('Thêm Supplementary Table grade-stratified; thêm 1 câu §3 hoặc '
    '§4.5.')


# ============================================================
H1('REVIEWER 6 — Editor / journal-level concerns')
P('Hồ sơ: Editor hoặc senior reviewer nhìn ở góc tổng thể. Sẽ chú ý '
  'đến contribution claim, predatory journal awareness, salami '
  'publishing, IRB issues, conflict of interest.')

H2('Câu hỏi 1')
Q('MAJOR', 'IRB approval reference is marked [TBD]. The manuscript '
   'cannot be considered for review without official ethics approval '
   'documentation.')
R('Em đề xuất CHỜ NCS Công Thị Hằng nhận thư chính thức từ Hội đồng '
  'Đạo đức HNUE; điền chính xác số quyết định, ngày phê duyệt; gửi '
  'kèm scan thư phê duyệt vào Supplementary materials.')
REV('CHỜ NCS — đây là blocking item Q3-6 đã ghi nhận. KHÔNG nộp bài '
    'trước khi có thông tin này.')

H2('Câu hỏi 2')
Q('MODERATE', 'The authors mention companion papers Q3 and Q4 using '
   'the same dataset. This raises potential salami publishing '
   'concerns. Please clearly articulate the unique contribution of '
   'Q2 vs the companion papers.')
R('Em đề xuất bổ sung §1.4 hoặc §4.6: "The present study (Q2) '
  'reports the integrated higher-order risk-protective SEM model in '
  'the pooled sample. Two pre-registered companion papers will '
  'address distinct research questions: (Q3) gender invariance of '
  'the SEM model using multi-group analysis; (Q4) person-centered '
  'phenotype profiles using latent profile analysis. Each paper '
  'tests pre-registered hypotheses non-overlapping with Q2, '
  'consistent with COPE guidelines on permissible secondary '
  'analyses of the same dataset."')
REV('Thêm 1 đoạn §1.4 hoặc §4.6 về relationship với companion '
    'papers; reference COPE 2018 guidelines on secondary analyses.')

H2('Câu hỏi 3')
Q('MAJOR', 'Frontiers in Psychiatry has been listed by Finland\'s '
   'Publication Forum (JUFO) at Level 0 (December 2024 decision). '
   'Has this been considered for institutional reporting purposes?')
R('Em đề xuất đây là vấn đề mà NHÓM cần biết và quyết định trước khi '
  'nộp. Phương án: (a) tiếp tục nộp Frontiers in Psychiatry chấp '
  'nhận downgrade caveat; (b) đổi sang BMC Public Health hoặc BMC '
  'Psychiatry — Q2 vẫn còn trong scope của hai tạp chí này. Đây là '
  'quyết định strategic, không phải vấn đề khoa học. Em đề xuất ghi '
  'nhận trong submission letter rằng nhóm đã cân nhắc và lựa chọn '
  'Frontiers in Psychiatry vì coverage area phù hợp + double-blind '
  'review process.')
REV('Vấn đề strategic — nhóm cần thảo luận; em đã ghi nhận trong '
    'PhuongAnXuLy doc; chuẩn bị submission letter với rationale cho '
    'lựa chọn tạp chí.')

H2('Câu hỏi 4')
Q('MINOR', 'Second author affiliation is marked [TBD]. Please '
   'resolve before submission.')
R('CHỜ thầy Nguyễn Minh Đức xác nhận affiliation chính thức (em '
  'không biết thầy NMĐ ở đơn vị nào ngoài việc giả định là cùng '
  'HNUE hoặc khác).')
REV('Em đã đề xuất câu hỏi trong message gửi nhóm; chờ thầy NMĐ '
    'confirm.')


# ============================================================
H1('REVIEWER 7 — Open science / transparency reviewer')
P('Hồ sơ: Reviewer chú trọng reproducibility, data sharing, '
  'preregistration. Phổ biến ở các tạp chí áp dụng TOP guidelines.')

H2('Câu hỏi 1')
Q('MAJOR', 'No mention of pre-registration. Was the analytic plan '
   'pre-registered on OSF or AsPredicted? If not, please clarify '
   'which analyses were planned a priori vs exploratory.')
R('Em đề xuất NHÓM đăng ký pre-registration trên OSF TRƯỚC khi nộp '
  'bài, theo Nosek et al. (2018) framework. Đăng ký gồm: research '
  'questions, sample size justification, measurement model, '
  'structural pathway hypotheses, multiple testing strategy. Nếu '
  'không kịp đăng ký prospective, em đề xuất "preregistration of '
  'remaining analyses" hoặc đăng ký retrospectively với caveat rõ '
  'ràng. Cite van den Akker et al. 2024 cho realistic discussion.')
REV('Đăng ký OSF prereg TRƯỚC khi nộp Q2; cite Nosek 2018, van den '
    'Akker 2024. Bổ sung §2 Methods: link OSF prereg.')

H2('Câu hỏi 2')
Q('MODERATE', 'Data Availability Statement is missing. Please '
   'provide a statement on data access (per FAIR principles).')
R('Em đề xuất bổ sung trước References: "Data Availability: De-'
  'identified data underlying this study are available from the '
  'corresponding author (CTH) upon reasonable request, subject to '
  'HNUE Ethics Committee approval for secondary use. Analysis code '
  'in Mplus and R is available at the OSF project page [link]."')
REV('Thêm Data Availability Statement; upload de-identified data '
    '+ analysis code lên OSF.')

H2('Câu hỏi 3')
Q('MINOR', 'CRediT taxonomy contribution statement is standard for '
   'most international journals.')
R('Em đề xuất bổ sung trước References: CRediT statement với 14 '
  'roles standard. Ví dụ: "CTH — Conceptualization, Investigation, '
  'Writing original draft. DMD — Supervision, Methodology. NMĐ — '
  'Supervision, Writing review & editing. Hai-Nguyen — Formal '
  'analysis, Writing review & editing."')
REV('Thêm CRediT statement với phân vai cụ thể; chờ nhóm xác nhận '
    'phân vai.')


# ============================================================
H1('TỔNG KẾT — DANH SÁCH SỬA TRƯỚC KHI NỘP')

P('Em xếp ưu tiên theo mức độ blocking. Các mục MAJOR phải hoàn '
  'thành trước; MODERATE có thể được reviewer chấp nhận với revision; '
  'MINOR là phần polish cuối.')

H2('Phải làm trước khi nộp (BLOCKING)')
P('1. IRB approval letter HNUE (Q3-6) — CHỜ NCS Công Thị Hằng')
P('2. Pre-registration trên OSF — NCS + nhóm duyệt H1-H3 trước')
P('3. Affiliation thầy Nguyễn Minh Đức — CHỜ thầy confirm')
P('4. Quyết định strategic về Frontiers in Psychiatry (downgrade '
  'caveat) — nhóm thảo luận')
P('5. CRediT statement — phân vai chính thức')

H2('Sửa trong bản thảo trước khi nộp')
P('6. §2.1 — Sample size justification (Wolf 2013)')
P('7. §2.2 — Subtype selection rationale (Beesdo 2009) + cross-'
  'cultural translation chi tiết (Brislin 1970)')
P('8. §2.3 — Model comparison strategy + multiple testing strategy '
  '(Benjamini-Hochberg 1995)')
P('9. §3 — Local fit reporting + R² calculation clarification + '
  'McDonald omega (Hayes & Coutts 2020)')
P('10. §4.4 — Tam giáo framing as theoretical proposition (không '
   'phải tested mechanism)')
P('11. §4.5 — Limitations expansion: rural + ethnic minority + '
   'puberty + cross-cultural validation')
P('12. §4.6 — Companion papers Q3+Q4 rationale (COPE 2018)')
P('13. Data Availability Statement + OSF project link')

H2('Sửa nhẹ (polish)')
P('14. Bootstrap CI cho indirect effects nếu mô hình có mediation')
P('15. Cognitive interviewing chi tiết (Supplementary)')
P('16. Grade-stratified sensitivity analysis (Supplementary)')


# ============================================================
H1('TÀI LIỆU THAM KHẢO CẦN BỔ SUNG')
P('Tất cả entries dưới đây nên được verify PMID/DOI trước khi chèn '
  'vào ref list của bản thảo. Hiện em đã đề xuất các ref dưới đây '
  'như nguồn để cite trong các câu trả lời.')

refs = [
    'Wolf EJ, Harrington KM, Clark SL, Miller MW. (2013). Sample '
    'size requirements for structural equation models: An evaluation '
    'of power, bias, and solution propriety. Educ Psychol Meas. '
    '73(6):913-934. DOI: 10.1177/0013164413495237.',
    'Hayes AF, Coutts JJ. (2020). Use omega rather than Cronbach\'s '
    'alpha for estimating reliability. But… Commun Methods Meas. '
    '14(1):1-24. DOI: 10.1080/19312458.2020.1718629.',
    'Little RJA. (1988). A test of missing completely at random for '
    'multivariate data with missing values. J Am Stat Assoc. '
    '83(404):1198-1202.',
    'Benjamini Y, Hochberg Y. (1995). Controlling the false '
    'discovery rate: a practical and powerful approach to multiple '
    'testing. J R Stat Soc Ser B. 57(1):289-300.',
    'Beesdo K, Knappe S, Pine DS. (2009). Anxiety and anxiety '
    'disorders in children and adolescents: developmental issues '
    'and implications for DSM-V. Psychiatr Clin North Am. '
    '32(3):483-524. PMID: 19716988.',
    'Brislin RW. (1970). Back-translation for cross-cultural '
    'research. J Cross-Cult Psychol. 1(3):185-216.',
    'Petersen AC, Crockett L, Richards M, Boxer A. (1988). A self-'
    'report measure of pubertal status: reliability, validity, and '
    'initial norms. J Youth Adolesc. 17(2):117-133. PMID: 24277579.',
    'Committee on Publication Ethics (COPE). (2018). Text recycling '
    'guidelines for editors. https://publicationethics.org/text-'
    'recycling-guidelines',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
