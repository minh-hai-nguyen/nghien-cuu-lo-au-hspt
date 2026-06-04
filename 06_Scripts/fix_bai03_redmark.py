# -*- coding: utf-8 -*-
"""
Bài 03: sửa nhãn Bảng 5 + BÔI ĐỎ mọi chỗ cần tác giả xử lý.
Đầu vào: Công Thị Hằng_v4_đã sửa.docx (đã có các fix văn bản).
Đầu ra:  Công Thị Hằng_v5_đánh dấu đỏ.docx
Đỏ = việc tác giả phải kiểm tra/sửa rồi xóa dấu đỏ trước khi nộp.
"""
import os
from docx import Document
from docx.shared import RGBColor

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/bài-03"
IN = os.path.join(ROOT, "Công Thị Hằng_v4_đã sửa.docx")
OUT = os.path.join(ROOT, "Công Thị Hằng_v5_đánh dấu đỏ.docx")
RED = RGBColor(0xFF, 0x00, 0x00)


def font_of(p):
    if p.runs:
        return p.runs[0].font.name, p.runs[0].font.size
    return None, None


def append_note(doc, anchor, note):
    for p in doc.paragraphs:
        if anchor in p.text:
            name, size = font_of(p)
            r = p.add_run("  【" + note + "】")
            r.font.color.rgb = RED
            r.bold = True
            if name:
                r.font.name = name
            if size:
                r.font.size = size
            return True
    return False


def redden_all(doc, target):
    """Bôi đỏ MỌI lần xuất hiện target trong các đoạn body."""
    n = 0
    for p in doc.paragraphs:
        if target in p.text:
            full = p.text
            name, size = font_of(p)
            idx = full.index(target)
            before, after = full[:idx], full[idx + len(target):]
            for r in list(p.runs):
                r.text = ""
            if p.runs:
                p.runs[0].text = before
            mid = p.add_run(target)
            mid.font.color.rgb = RED
            mid.bold = True
            aft = p.add_run(after)
            for rr in (mid, aft):
                if name:
                    rr.font.name = name
                if size:
                    rr.font.size = size
            n += 1
    return n


def redden_cell(cell, new_text):
    p = cell.paragraphs[0]
    name, size = font_of(p)
    for r in list(p.runs):
        r.text = ""
    if p.runs:
        p.runs[0].text = new_text
        run = p.runs[0]
    else:
        run = p.add_run(new_text)
    run.font.color.rgb = RED
    run.bold = True
    if name:
        run.font.name = name
    if size:
        run.font.size = size


doc = Document(IN)
log = []

# --- 1. Bảng 5 (d.tables[4]) — sửa nhãn 3 mệnh đề + bôi đỏ ---
tb5 = doc.tables[4]
labels = ["Mối quan hệ của em với cha",
          "Mối quan hệ của em với mẹ",
          "Mối quan hệ của em với anh/chị/em"]
for i, lab in enumerate(labels):
    redden_cell(tb5.rows[i + 2].cells[1], lab)
log.append(("Bảng 5: sửa+đỏ 3 nhãn mệnh đề", True))
log.append(("Bảng 5: ghi chú", append_note(doc, "Bảng 5. Quan hệ gia đình",
    "ĐÃ SỬA NHÃN MỆNH ĐỀ cho khớp số liệu và đoạn diễn giải — tác giả đối chiếu "
    "lại với bảng hỏi gốc rồi xóa dấu đỏ")))

# --- 2. Bảng 8 — ghi chú cần dựng lại ---
log.append(("Bảng 8: ghi chú", append_note(doc, "Bảng 8.",
    "CẦN DỰNG LẠI BẢNG: cột 'Sai số chuẩn của ước lượng' đang bị lặp 2 lần và "
    "thiếu cột R² chưa hiệu chỉnh; giá trị Durbin-Watson 2,010 bị lặp. Lấy lại từ "
    "kết xuất SPSS — cấu trúc đúng: Mô hình | R | R² | R² hiệu chỉnh | Sai số chuẩn "
    "của ước lượng | Durbin-Watson")))

# --- 3. ANOVA — bôi đỏ + ghi chú ---
log.append(("ANOVA: bôi đỏ", redden_all(doc,
    "phân tích phương sai một yếu tố (one-way ANOVA)") > 0))
log.append(("ANOVA: ghi chú", append_note(doc,
    "Dữ liệu được xử lý và phân tích bằng phần mềm SPSS",
    "CẦN SỬA: không thấy kết quả ANOVA trong mục Kết quả — bổ sung kết quả phân "
    "tích phương sai, hoặc bỏ phương pháp này khỏi danh sách")))

# --- 4. Khoảng tổng điểm thang lo âu ---
log.append(("Tổng điểm: ghi chú", append_note(doc, "Bảng 6. Rối loạn lo âu",
    "CẦN BỔ SUNG: nêu rõ khoảng tổng điểm thang lo âu (5 mục × thang 1-4 = 5-20); "
    "đoạn diễn giải đang gọi tổng điểm 10,95 là 'điểm trung bình' — cần sửa thành "
    "'tổng điểm' hoặc 'điểm trung bình tổng thang đo'")))

# --- 5. Địa bàn: tại Việt Nam vs tại Hà Nội ---
nloc = redden_all(doc, "tại Việt Nam")
log.append((f"Địa bàn: bôi đỏ 'tại Việt Nam' ({nloc})", nloc > 0))
log.append(("Địa bàn: ghi chú", append_note(doc,
    "Mẫu khách thể trong nghiên cứu này bao gồm 433",
    "THỐNG NHẤT ĐỊA BÀN: bài lúc ghi 'tại Việt Nam' lúc 'tại Hà Nội' (abstract "
    "tiếng Anh ghi Hanoi) — chọn một địa bàn và dùng nhất quán toàn bài")))

# --- 6. Placeholder thông tin tác giả ---
for ph in ["[Tác giả tự xác nhận — ví dụ: Nghiên cứu sinh / Thạc sĩ / Tiến sĩ]",
           "[Địa chỉ đầy đủ của đơn vị — tác giả bổ sung]",
           "[Tác giả bổ sung nếu có — đăng ký miễn phí tại https://orcid.org]"]:
    log.append((f"Placeholder đỏ: {ph[:30]}...", redden_all(doc, ph) > 0))

# --- 7. Công cụ đo APA DSM-5 ---
log.append(("Công cụ đo: ghi chú", append_note(doc,
    "(1) Thang đo rối loạn lo âu tổng quát, gồm 5 mục",
    "CẦN ĐỐI CHIẾU CÔNG CỤ: thang lo âu dùng 5 mục/thang 1-4, nhưng tài liệu tham "
    "khảo trích thang APA DSM-5 (bản gốc 10 mục, thang 0-4). Ghi rõ đây là bản rút "
    "gọn/thích nghi, hoặc trích đúng nguồn của thang 5 mục đã sử dụng")))

doc.save(OUT)
print("Kết quả đánh dấu Bài 03:")
for k, v in log:
    print(f"  {'OK ' if v else '!! '}{k}")
print(f"\nĐã lưu: {OUT}")
