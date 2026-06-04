"""DOC: Kiem tra chenh lech gioi RLLAXH vs RLLATQ — 3 chi so
ALL FACTS FROM Bang 3.20 chuong 3 luan an, file v2:
- n nam=614, n nu=738
- RLLATQ: nam M=51,43 SD=22,010; nu M=59,47 SD=22,072
- RLLAXH: nam M=43,20 SD=25,093; nu M=52,74 SD=26,311
- F (gioi x RLLATQ) = 44,484
- F (gioi x RLLAXH) = 45,984

PHEP TOAN COHEN d:
- d RLLATQ = 0,3647
- d RLLAXH = 0,3703
- Ty so: 1,015 (XH lon hon 1,52%)
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Kiem_tra_chenh_lech_gioi_RLLAXH_vs_RLLATQ.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('KIỂM TRA CHÊNH LỆCH GIỚI TÍNH\nRLLAXH vs RLLATQ — 3 chỉ số khác nhau\n— Phép toán chi tiết từ Bảng 3.20 chương 3 —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Em kiểm tra lại số liệu thực trạng xem ở lo âu xã hội, '
    'chênh lệch giới tính có cao hơn lo âu tổng quát không nhé?',
    italic=True
)

# Trả lời ngắn
H('Trả lời ngắn — phụ thuộc CHỈ SỐ', level=2, color=NAVY)
para(
    'Câu trả lời PHỤ THUỘC vào chỉ số sử dụng. Có 3 cách đo "chênh '
    'lệch giới" khác nhau, cho kết quả CHÊNH LỆCH NHAU đáng kể:', bold=True
)
add_table(
    ['Chỉ số', 'RLLATQ (lan tỏa)', 'RLLAXH (xã hội)', 'XH lớn hơn TQ', 'Đánh giá'],
    [
        ['Chênh lệch ĐTB raw (nữ − nam)', '8,04 điểm', '9,54 điểm', '+18,7%', 'Khác biệt RÕ RỆT'],
        ['F-test (giới × loại)', '44,484', '45,984', '+3,37%', 'Khác biệt NHỎ'],
        ['Cohen d (effect size chuẩn)', '0,3647', '0,3703', '+1,52%', 'GẦN NHƯ BẰNG'],
    ]
)
para('')
para(
    'Như vậy, KẾT QUẢ thay đổi tùy chỉ số: nếu nhìn chênh lệch ĐTB '
    'thô, RLLAXH lớn hơn RLLATQ 18,7% — RỌ RỆT. Nếu nhìn theo '
    'effect size chuẩn (Cohen d), HAI giá trị GẦN NHƯ BẰNG NHAU '
    '(0,370 vs 0,365 — chênh chỉ 1,5%).'
)

# I. Số liệu gốc Bảng 3.20
H('I. Số liệu gốc từ Bảng 3.20 chương 3', level=2, color=NAVY)
add_table(
    ['Biến', 'Nhóm', 'n', 'ĐTB', 'ĐLC (SD)'],
    [
        ['RLLATQ (lan tỏa)', 'Nam', '614', '51,43', '22,010'],
        ['RLLATQ (lan tỏa)', 'Nữ', '738', '59,47', '22,072'],
        ['RLLAXH (xã hội)', 'Nam', '614', '43,20', '25,093'],
        ['RLLAXH (xã hội)', 'Nữ', '738', '52,74', '26,311'],
    ]
)
para('Tổng mẫu: n = 1.352 (614 nam + 738 nữ).', italic=True)

# II. Cách 1 — Chênh lệch ĐTB
H('II. Cách 1 — Chênh lệch ĐTB (raw)', level=2, color=NAVY)
para('Phép toán đơn giản: ĐTB nữ − ĐTB nam.')
para('• RLLATQ: 59,47 − 51,43 = 8,04 điểm')
para('• RLLAXH: 52,74 − 43,20 = 9,54 điểm')
para('Tỷ số: 9,54 / 8,04 = 1,1866 → RLLAXH lớn hơn RLLATQ 18,7%.', bold=True)
para(
    'Hạn chế: chỉ số này KHÔNG xét đến phương sai (variance) trong '
    'mỗi nhóm. RLLAXH có SD lớn hơn RLLATQ → chênh lệch raw có thể '
    'bị "thổi phồng" nếu chỉ xem ĐTB.'
)

# III. Cách 2 — F-test
H('III. Cách 2 — F-test (One-way ANOVA)', level=2, color=NAVY)
para('Tác giả luận án đã báo cáo F-test trong Bảng 3.20:')
para('• F (giới × RLLATQ) = 44,484; p < 0,001')
para('• F (giới × RLLAXH) = 45,984; p < 0,001')
para('Tỷ số: 45,984 / 44,484 = 1,0337 → RLLAXH lớn hơn 3,37%.')
para(
    'Hạn chế: F-test phụ thuộc vào CỠ MẪU và VARIANCE. Với n = '
    '1.352 cỡ mẫu lớn, F-test rất nhạy — khác biệt nhỏ về ĐTB '
    'đã có thể tạo F-test cao. F-test phù hợp cho KIỂM ĐỊNH ý '
    'nghĩa thống kê, nhưng KHÔNG đo CƯỜNG ĐỘ HIỆU ỨNG.'
)
para(
    'Lưu ý: F = (MS giữa nhóm) / (MS trong nhóm). Khi VARIANCE '
    'tổng (SD pooled) tăng, MS trong nhóm cũng tăng — F có thể '
    'KHÔNG tăng tương ứng với chênh ĐTB. Đó là lý do F RLLAXH '
    'tăng chỉ 3,37% trong khi chênh ĐTB tăng 18,7%.'
)

# IV. Cách 3 — Cohen d
H('IV. Cách 3 — Cohen d (effect size chuẩn hóa)', level=2, color=NAVY)
para('Công thức: d = (M_nữ − M_nam) / SD_pooled')
para('Trong đó: SD_pooled = √[((n₁−1)·SD₁² + (n₂−1)·SD₂²) / (n₁+n₂−2)]')

para('')
para('RLLATQ:', bold=True)
para('• SD pooled = √[(613·22,010² + 737·22,072²) / 1350]')
para('             = √[(296.962 + 358.927) / 1350]')
para('             = √485,84 = 22,044')
para('• d = 8,04 / 22,044 = 0,3647', bold=True)

para('')
para('RLLAXH:', bold=True)
para('• SD pooled = √[(613·25,093² + 737·26,311²) / 1350]')
para('             = √[(385.992 + 510.203) / 1350]')
para('             = √663,85 = 25,765')
para('• d = 9,54 / 25,765 = 0,3703', bold=True)

para('')
para('Tỷ số d: 0,3703 / 0,3647 = 1,0152 → RLLAXH lớn hơn 1,52%.', bold=True)

para('')
para(
    'Theo Cohen (1988): d = 0,2 (small), 0,5 (medium), 0,8 (large). '
    'Cả d RLLATQ (0,365) và d RLLAXH (0,370) đều thuộc nhóm '
    'SMALL-TO-MEDIUM, gần ngưỡng SMALL. Hai effect size GẦN NHƯ '
    'BẰNG NHAU.'
)

# V. Lý do 3 chỉ số khác nhau
H('V. Vì sao 3 chỉ số cho kết quả khác nhau?', level=2, color=NAVY)
para(
    'Khác biệt chính nằm ở ĐỘ LỆCH CHUẨN (variance):'
)
para('• SD pooled RLLATQ = 22,044')
para('• SD pooled RLLAXH = 25,765')
para('Tỷ số: 25,765 / 22,044 = 1,1689 — SD RLLAXH lớn hơn 16,9%.')
para('')
para(
    'RLLAXH có biến thiên LỚN HƠN giữa các học sinh (SD = 25,8 vs '
    '22,0). Khi chuẩn hóa theo SD (Cohen d), chênh lệch RAW 9,54 '
    '"co lại" tương đối — kết quả d gần như bằng RLLATQ.'
)
para(
    'Diễn giải đơn vị: cả nam và nữ đều có biến thiên LỚN HƠN '
    'trong mức độ lo âu xã hội so với lo âu lan tỏa. Có thể do '
    'lo âu xã hội PHÂN HÓA mạnh hơn — một số học sinh rất sợ '
    'tình huống xã hội, một số rất bình thản. Lo âu lan tỏa '
    'phân bố ĐỀU HƠN.'
)

# VI. Câu trả lời chính xác
H('VI. Câu trả lời CHÍNH XÁC', level=2, color=NAVY)
blue_run('Trả lời câu hỏi của thầy:', bold=True)
blue_run(
    'CÓ — chênh lệch giới ở RLLAXH cao hơn RLLATQ, NHƯNG MỨC ĐỘ '
    'phụ thuộc chỉ số:'
)
blue_run(
    '(1) THEO CHÊNH LỆCH ĐTB RAW: RLLAXH (9,54 điểm) cao hơn '
    'RLLATQ (8,04 điểm) — chênh 18,7%. Đây là khác biệt RỌ RỆT '
    'trên thang điểm thực.'
)
blue_run(
    '(2) THEO F-TEST: RLLAXH (45,984) cao hơn RLLATQ (44,484) — '
    'chênh chỉ 3,4%. F-test bị "kéo xuống" do RLLAXH có '
    'variance lớn hơn.'
)
blue_run(
    '(3) THEO COHEN d (chuẩn hóa): RLLAXH (0,370) cao hơn RLLATQ '
    '(0,365) — chênh CHỈ 1,5%. Hai effect size GẦN NHƯ BẰNG '
    'NHAU. Cả hai đều thuộc nhóm small-to-medium (Cohen 1988).'
)
blue_run(
    '(4) KIẾN NGHỊ CHO BÁO CÁO CTH: KHÔNG nên nói "RLLAXH chênh '
    'lệch giới RỌ RỆT NHẤT" mà không nêu rõ chỉ số. Báo cáo '
    'chính xác:\n'
    '   • F-test: RLLAXH (45,984) > RLLATQ (44,484), p < 0,001\n'
    '   • Chênh lệch ĐTB: RLLAXH 9,54 vs RLLATQ 8,04 điểm\n'
    '   • Cohen d: 0,370 (RLLAXH) ≈ 0,365 (RLLATQ) — small-to-medium\n'
    '   • Cả hai đều có ý nghĩa thống kê cao (p < 0,001)\n'
    '   • Mức độ vượt trội của RLLAXH so với RLLATQ phụ thuộc '
    'chỉ số đo'
)

# VII. Sửa chữa các doc trước
H('VII. Sửa chữa các doc đã viết trước đây', level=2, color=NAVY)
para(
    'Trong các doc trước (Doc 4 H4 gender + Doc BNHĐ + AUDIT), em '
    'có viết "F RLLAXH (45,984) > F RLLATQ (44,484) — chênh lệch '
    'giới ở lo âu xã hội RỌ RỆT NHẤT". Phát biểu này CHÍNH XÁC '
    'theo F-test nhưng có thể GÂY HIỂU LẦM về CƯỜNG ĐỘ HIỆU ỨNG.', bold=True, color=RED
)
para('Đề xuất sửa lại trong các doc đó:')
para(
    '• Thay "F RLLAXH > F RLLATQ → chênh lệch giới ở RLLAXH RỌ RỆT '
    'NHẤT" → "F RLLAXH > F RLLATQ về ý nghĩa thống kê. Tuy nhiên '
    'effect size (Cohen d) gần như bằng nhau (0,370 vs 0,365), gợi '
    'ý CƯỜNG ĐỘ thực tế tương đương."', italic=True
)

# VIII. So sánh với loại RLLAC để hoàn thiện bức tranh
H('VIII. So sánh với RLLAC (lo âu chia ly) — hoàn thiện bức tranh', level=2, color=NAVY)
para('Tính d cho RLLAC để có bức tranh đầy đủ ba dạng RLLA:')
add_table(
    ['Biến', 'Nam ĐTB (SD)', 'Nữ ĐTB (SD)', 'Chênh ĐTB', 'F', 'p', 'Cohen d'],
    [
        ['RLLATQ', '51,43 (22,010)', '59,47 (22,072)', '+8,04', '44,484', '< 0,001', '0,365'],
        ['RLLAXH', '43,20 (25,093)', '52,74 (26,311)', '+9,54', '45,984', '< 0,001', '0,370'],
        ['RLLAC',  '25,42 (25,459)', '24,76 (23,294)', '−0,66', '0,246',  '0,620',   '~0,03 (NS)'],
    ]
)
para('')
para(
    'Bức tranh đầy đủ: chênh lệch giới ở RLLATQ và RLLAXH GẦN '
    'NHƯ BẰNG NHAU (d ≈ 0,37); RLLAC KHÔNG có chênh lệch giới '
    '(d ≈ 0,03, NS). Pattern phù hợp giả thuyết H4 đã phát biểu.'
)

# Kết luận
H('IX. Kết luận', level=2, color=NAVY)
p = d.add_paragraph()
r = p.add_run(
    'Cảm ơn thầy đã yêu cầu kiểm tra lại — phát hiện này giúp '
    'em diễn giải CHÍNH XÁC HƠN trong các báo cáo CTH. Tóm '
    'lại: chênh lệch giới ở RLLAXH cao hơn RLLATQ theo CHÊNH '
    'ĐTB (18,7%) và F-test (3,4%) — nhưng theo Cohen d '
    '(effect size chuẩn hóa), HAI giá trị GẦN NHƯ BẰNG NHAU '
    '(0,370 vs 0,365 — chênh chỉ 1,5%). RLLAXH có variance '
    'lớn hơn (SD ≈ 25,8) nên d thấp hơn so với mức chênh '
    'ĐTB raw. Kiến nghị báo cáo CTH nêu rõ CẢ ba chỉ số để '
    'tránh hiểu lầm.'
)
r.font.size = Pt(12); r.font.color.rgb = BLUE; r.italic = True

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
