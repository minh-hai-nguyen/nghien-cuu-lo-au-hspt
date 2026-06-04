"""Bản dịch QT065 Bradshaw/Lochman 2025 EACP — full bilingual key sections."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/03_Ban-dich/QT065_Bradshaw_Lochman_EACP_RCT_JSchPsy_2025.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
DARK = RGBColor(31, 73, 125); GREEN = RGBColor(54, 95, 44); RED = RGBColor(192, 0, 0); GRAY = RGBColor(90, 90, 90); ORANGE = RGBColor(191, 97, 14)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def en(text):
    p = d.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
def vn(text, bold=False):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(11)
def page(n):
    p = d.add_paragraph(); r = p.add_run(f'--- Trang {n} ---'); r.bold = True; r.font.color.rgb = ORANGE; r.font.size = Pt(10)

# COVER
t = d.add_heading('QT065 — Bản dịch chi tiết: Early Adolescent Coping Power RCT (Bradshaw, Lochman et al. 2025)', level=0)
for r in t.runs: r.font.color.rgb = DARK
sub = d.add_paragraph(); sr = sub.add_run('Bản dịch song ngữ Anh-Việt — focus Abstract, Introduction, Methods, Results, Discussion. References giữ tiếng Anh trong PDF gốc.')
sr.italic = True; sr.font.size = Pt(11); sr.font.color.rgb = GRAY
d.add_paragraph()

# THÔNG TIN
add_h('Thông tin thư mục', 1)
meta = d.add_table(rows=10, cols=2); meta.style = 'Table Grid'
mrows = [
    ('Tiêu đề (EN)', 'Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers'),
    ('Tiêu đề (VN)', 'Thử nghiệm RCT Chương trình "Sức mạnh Ứng phó" cho VTN sớm: Hiệu quả lên các vấn đề cảm xúc và hành vi của học sinh THCS'),
    ('Tác giả', 'Bradshaw CP, McDaniel H, Pas ET, Debnam KJ, Bottiani JH, Powell N, Ialongo NS, Morgan-Lopez A, Lochman JE'),
    ('Email tác giả liên hệ', 'cpb8g@virginia.edu (Catherine P. Bradshaw, UVA)'),
    ('Đơn vị', 'University of Virginia + University of South Carolina + Johns Hopkins Bloomberg School + University of Alabama + RTI International'),
    ('Tạp chí', 'Journal of School Psychology, vol. 110, June 2025, art. 101437'),
    ('DOI', '10.1016/j.jsp.2025.101437'),
    ('PubMed', 'https://pubmed.ncbi.nlm.nih.gov/40506167/'),
    ('Ngày nhận / chấp nhận', 'Received 20 January 2024 | Revised 28 February 2025 | Accepted 28 February 2025'),
    ('Số trang / từ khoá', '15 trang. Keywords: Prevention; Aggressive behavior; RCT; Middle school; Indicated programming'),
]
for i, (k, v) in enumerate(mrows):
    c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
    meta.rows[i].cells[1].text = v
d.add_paragraph()

# ============ ABSTRACT ============
page(1)
add_h('TÓM TẮT (Abstract)', 1)

en('We report findings from a 40 middle school randomized controlled trial of an adapted version of Coping Power (Lochman & Wells, 2002a) for middle schoolers, called the Early Adolescent Coping Power (EACP) Program (Bradshaw et al., 2019) to determine the impact of EACP on adolescents\' mental health outcomes, as indicated by student self-reported and teacher-ratings on the Behavior Assessment System for Children (BASC).')
vn('Chúng tôi báo cáo kết quả từ một THỬ NGHIỆM RCT trên 40 trường THCS về phiên bản thích ứng của Coping Power (Lochman & Wells, 2002a) dành cho HS THCS, gọi là Chương trình EACP — Early Adolescent Coping Power (Bradshaw và cộng sự, 2019) — nhằm xác định tác động của EACP lên KẾT QUẢ SỨC KHOẺ TÂM THẦN của vị thành niên, đo bằng tự báo cáo của HS và đánh giá của giáo viên trên Hệ thống Đánh giá Hành vi cho Trẻ em (BASC).')
d.add_paragraph()

en('The EACP was implemented over the course of students\' 7th grade school year. The sample included 709 students who were identified at baseline through a teacher screening process for aggressive behavior and enrolled into the project (69.8% African American and 59.4% male).')
vn('EACP được triển khai trong suốt NĂM HỌC LỚP 7. Mẫu gồm 709 HS được xác định tại baseline qua quy trình SÀNG LỌC CỦA GIÁO VIÊN về hành vi gây hấn và đăng ký vào dự án (69,8 % người Mỹ gốc Phi, 59,4 % nam).')
d.add_paragraph()

en('For teacher-reported outcomes, intent to treat (ITT) results indicated that EACP was associated with a decrease in externalizing problems over time. Exploration of effect modification suggested that girls in the EACP condition demonstrated statistically significant decreases in school problems over time, as well as a baseline by intervention effect whereby students and schools with greater baseline difficulties demonstrated decreased internalizing problems over time.')
vn('Đối với KẾT QUẢ THEO BÁO CÁO CỦA GIÁO VIÊN, phân tích intent-to-treat (ITT) cho thấy EACP có liên quan đến GIẢM các VẤN ĐỀ NGOẠI HOÁ (externalizing problems) theo thời gian. Khám phá hiệu ứng điều biến cho thấy NỮ trong điều kiện EACP có giảm CÓ Ý NGHĨA THỐNG KÊ các vấn đề học đường theo thời gian, cũng như một hiệu ứng baseline-by-intervention: HS và trường có khó khăn baseline cao GIẢM được các vấn đề NỘI HOÁ (internalizing) theo thời gian.')
d.add_paragraph()

en('For student-reported outcomes, there was a significant moderated effect, whereby girls in EACP demonstrated more favorable self-reported personal adjustment outcomes. Together, these results suggest that the early adolescent adaptation of Coping Power had preventive effects for 7th grade participants, and promotive effects specifically for girls, over time.')
vn('Đối với KẾT QUẢ TỰ BÁO CÁO CỦA HS, có một hiệu ứng điều biến có ý nghĩa thống kê: NỮ trong EACP cho thấy kết quả TỰ ĐIỀU CHỈNH CÁ NHÂN (personal adjustment) tốt hơn. Tổng hợp lại, những kết quả này gợi ý rằng phiên bản EACP của Coping Power có hiệu ứng PHÒNG NGỪA cho HS lớp 7, và hiệu ứng THÚC ĐẨY (promotive) đặc biệt cho nữ giới, theo thời gian.')
d.add_paragraph()

# ============ INTRO ============
add_h('1. GIỚI THIỆU (Introduction) — Các điểm chính', 1)

en('Aggressive and disruptive behavior problems often co-occur with poor academic functioning and low school connectedness. These issues are particularly salient in middle school, where rates of school violence and disruption are higher relative to elementary school and student engagement and caregiver involvement in school-based programming decline.')
vn('Các vấn đề HÀNH VI GÂY HẤN và GÂY RỐI thường ĐỒNG XUẤT HIỆN với chức năng học tập kém và sự kết nối với trường thấp. Những vấn đề này đặc biệt nổi bật ở THCS, nơi tỷ lệ bạo lực và gây rối trường học cao hơn TIỂU HỌC, và sự tham gia của HS cùng phụ huynh giảm.')
d.add_paragraph()

en('Yet there are fewer school-based prevention programs to address the behavioral and mental health concerns of middle school students, especially programs for students engaging in high rates of aggressive behaviors and demonstrating other indicators of risk for disengagement from school, academic problems, and longer-term mental health concerns.')
vn('Tuy nhiên, có ÍT chương trình phòng ngừa dựa vào trường để giải quyết các vấn đề hành vi và sức khoẻ tâm thần của HS THCS, đặc biệt là các chương trình cho HS có tỷ lệ hành vi gây hấn cao và thể hiện các dấu hiệu nguy cơ tách rời khỏi trường, vấn đề học tập, và lo ngại sức khoẻ tâm thần dài hạn.')
d.add_paragraph()

en('Coping Power (Lochman & Wells, 2002a) is a rigorously tested, school-based program that aims to prevent aggressive behavior problems in upper elementary students. This multi-component program includes clinician-facilitated group sessions for youth, separate group sessions for caregivers (i.e., parents), and select program activities for teachers.')
vn('Coping Power (Lochman & Wells, 2002a) là một chương trình DỰA VÀO TRƯỜNG được kiểm thử nghiêm ngặt nhằm phòng ngừa các vấn đề hành vi gây hấn ở HS tiểu học cấp cao. Chương trình ĐA THÀNH PHẦN này gồm các buổi nhóm dành cho TRẺ do nhà chuyên môn lâm sàng điều phối, các buổi nhóm RIÊNG dành cho NGƯỜI CHĂM SÓC (cha mẹ), và các hoạt động chọn lọc cho GIÁO VIÊN.')
d.add_paragraph()

en('Though found to be effective with students in the upper elementary grades, prior to the present study, Coping Power had rarely been tested with youth in the middle school years, beyond 6th grade. Given that previous studies have demonstrated a range of positive effects for Coping Power with upper elementary school children (i.e., grades 4–6, within the elementary school setting), there was potential that these effects could also generalize to students in middle school.')
vn('Mặc dù được chứng minh hiệu quả với HS các lớp cuối tiểu học, TRƯỚC nghiên cứu hiện tại, Coping Power HIẾM khi được kiểm thử với HS THCS (từ lớp 7 trở lên). Do các nghiên cứu trước đã cho thấy nhiều hiệu ứng tích cực của Coping Power với trẻ tiểu học cấp cao (lớp 4-6, trong bối cảnh tiểu học), có TIỀM NĂNG các hiệu ứng này có thể NGOẠI SUY cho HS THCS.')
d.add_paragraph()

en('However, to increase the developmental fit of select program elements and activities, we made systematic developmental adaptations to the intervention, in turn creating a slightly adapted version of the model called the Early Adolescent Coping Power (EACP) Program. Like the original, this adapted version of Coping Power was intended to meet the social, cognitive, and developmental needs of middle school students (grades 6–8) showing signs of difficulty self-regulating anger.')
vn('Tuy nhiên, để TĂNG TÍNH PHÙ HỢP PHÁT TRIỂN của các yếu tố và hoạt động chương trình được chọn, chúng tôi đã thực hiện các THÍCH ỨNG PHÁT TRIỂN HỆ THỐNG, từ đó tạo ra phiên bản hơi điều chỉnh gọi là Chương trình EACP. Giống bản gốc, phiên bản EACP này nhằm đáp ứng nhu cầu xã hội, nhận thức và phát triển của HS THCS (lớp 6-8) có dấu hiệu khó tự điều chỉnh CƠN GIẬN.')
d.add_paragraph()

# ============ EACP COMPONENTS ============
page(3)
add_h('2. MÔ TẢ EACP — 3 thành phần', 1)

add_h('2.1. Thành phần HỌC SINH (Youth Component)', 2)
en('The EACP Youth Component as delivered in the context of this randomized controlled trial (RCT) included both a group and individual component. The youth group component included 6–7 students and took place at the students\' schools, typically during nonacademic homeroom periods, recess, specials, or lunch. The 25 group sessions lasted approximately 45 min and were designed to fit within a single course period.')
vn('Thành phần HS của EACP triển khai trong RCT này gồm cả phần NHÓM và phần CÁ NHÂN. Phần nhóm HS gồm 6-7 HS và diễn ra TẠI TRƯỜNG của HS, thường vào tiết homeroom không-học-thuật, giờ ra chơi, lớp đặc biệt, hoặc giờ ăn trưa. 25 BUỔI NHÓM kéo dài khoảng 45 phút mỗi buổi, được thiết kế phù hợp 1 tiết học.')
d.add_paragraph()

en('Group youth sessions delivered in this RCT were held approximately once week over the school year (spanning nearly 9 months). The group sessions in this RCT were co-led by a grant-funded intervention staff member (i.e., master\'s, doctorate-level clinician, or clinician trainee) and when possible, were co-facilitated by the schools\' counselor, social worker, or school psychologist.')
vn('Các buổi nhóm HS trong RCT này được tổ chức KHOẢNG MỖI TUẦN MỘT LẦN trong suốt năm học (kéo dài gần 9 tháng). Các buổi nhóm trong RCT này được ĐỒNG ĐIỀU PHỐI bởi một nhân viên can thiệp được tài trợ (thạc sĩ, tiến sĩ lâm sàng, hoặc clinician trainee) và khi có thể, được đồng điều phối bởi cố vấn học đường, nhân viên xã hội, hoặc nhà tâm lý học đường.')
d.add_paragraph()

en('The first eight sessions (Phase 1) focused on the delivery of content and concepts to the students and on developing group norms, setting short- and long-term goals, organizational and study skills, awareness of feelings, anger coping and self-control, coping self-statements, relaxation techniques, and perspective-taking.')
vn('GIAI ĐOẠN 1 (8 buổi đầu) tập trung vào cung cấp NỘI DUNG và KHÁI NIỆM cho HS, và phát triển các quy tắc nhóm, đặt mục tiêu NGẮN HẠN + DÀI HẠN, kỹ năng TỔ CHỨC + HỌC TẬP, NHẬN THỨC VỀ CẢM XÚC, ỨNG PHÓ CƠN GIẬN + tự kiểm soát, các CÂU TỰ NÓI ĐỂ ỨNG PHÓ (coping self-statements), kỹ thuật THƯ GIÃN, và LẤY GÓC NHÌN (perspective-taking).')
d.add_paragraph()

en('In the 9th through 15th session (Phase 2), students learned to apply the concepts taught by practicing perspective-taking and social problem-solving using the "PICC" (Problem Identification Choices Consequences) model. This phase culminated in the student group planning and creating a video to teach others the concepts that they had learned.')
vn('GIAI ĐOẠN 2 (buổi 9-15), HS học ÁP DỤNG các khái niệm bằng cách thực hành lấy góc nhìn và giải quyết vấn đề xã hội qua mô hình "PICC" — Problem Identification, Choices, Consequences (Xác định Vấn đề, Lựa chọn, Hậu quả). Giai đoạn này kết thúc bằng việc nhóm HS LẬP KẾ HOẠCH và TẠO VIDEO để dạy lại cho người khác các khái niệm đã học.')
d.add_paragraph()

en('In the final phase (i.e., sessions 16–25), developmental adaptations were made to support the students in applying their skills to scenarios and challenges early adolescents face, including friendship and romantic relationship development, social aggression, cyberbullying, problem-solving about damaged relationships and neighborhood problems, refusal skills/coping with peer pressure, deviant peer groups and centrality of group membership, positive quality development, entering new peer groups, and using positive peer networks. Students were also taught to utilize assertive "I-statements."')
vn('GIAI ĐOẠN 3 (buổi 16-25), các thích ứng phát triển được thực hiện để hỗ trợ HS ÁP DỤNG kỹ năng vào các tình huống và thách thức mà VTN sớm gặp phải, gồm: phát triển TÌNH BẠN và quan hệ TÌNH CẢM, hành vi GÂY HẤN XÃ HỘI, BẮT NẠT QUA MẠNG (cyberbullying), giải quyết vấn đề về quan hệ HỎNG và vấn đề khu phố, kỹ năng TỪ CHỐI / ứng phó áp lực bạn bè, nhóm bạn LỆCH LẠC và tầm quan trọng của thành viên nhóm, phát triển CHẤT LƯỢNG TÍCH CỰC, hoà nhập NHÓM BẠN MỚI, và sử dụng mạng lưới bạn TÍCH CỰC. HS cũng được dạy sử dụng "I-statements" KHẲNG ĐỊNH.')
d.add_paragraph()

en('In addition, EACP was designed so that each child also received a total of 8–10 individual 30-min sessions with the group leaders during school hours to further build rapport, monitor progress, and reinforce content. The individual sessions were used primarily for monitoring and reinforcing attainment of classroom and social behavior goals (e.g., avoiding fights, resisting peer pressure) and coping with specific attributional biases and social problem-solving challenges.')
vn('Ngoài ra, EACP được thiết kế để mỗi HS còn nhận TỔNG 8-10 BUỔI CÁ NHÂN 30 phút với các trưởng nhóm trong giờ học nhằm xây dựng quan hệ, theo dõi tiến độ, và củng cố nội dung. Các buổi cá nhân chủ yếu để giám sát và củng cố việc đạt được mục tiêu hành vi lớp học và xã hội (vd: tránh đánh nhau, chống áp lực bạn bè) và đối phó với các thiên lệch quy gán cụ thể (attributional biases) và thách thức giải quyết vấn đề xã hội.')
d.add_paragraph()

add_h('2.2. Thành phần CARGIVER (Người chăm sóc / Cha mẹ)', 2)
en('The EACP Caregiver Component consisted of 12 group sessions for caregivers, implemented over the same school year as the student component. These sessions focused on building caregiver skills in behavior management, supporing the new skills students were learning in student groups, and to promote a more positive school experience for students. Caregivers met in groups of 10–12 with the two co-leaders of the youth component.')
vn('Thành phần CHĂM SÓC của EACP gồm 12 BUỔI NHÓM cho người chăm sóc, triển khai trong cùng năm học với thành phần HS. Các buổi này tập trung xây dựng KỸ NĂNG QUẢN LÝ HÀNH VI cho người chăm sóc, hỗ trợ kỹ năng mới mà HS đang học trong nhóm HS, và thúc đẩy trải nghiệm trường học TÍCH CỰC HƠN cho HS. Người chăm sóc gặp theo nhóm 10-12 người với 2 trưởng nhóm của thành phần HS.')
d.add_paragraph()

en('Over the course of the 12 sessions, caregivers received training on a variety of skills including rewarding wanted behaviors, giving effective instructions, establishing age-appropriate rules and expectations for adolescents, applying effective consequences to unwanted behavior, establishing ongoing family communication structures (e.g., weekly family meetings), and effective monitoring.')
vn('Trong 12 buổi, người chăm sóc được đào tạo nhiều kỹ năng: KHEN THƯỞNG hành vi mong muốn, ĐƯA HƯỚNG DẪN HIỆU QUẢ, thiết lập QUY TẮC và KỲ VỌNG phù hợp tuổi cho VTN, áp dụng HẬU QUẢ hiệu quả với hành vi không mong muốn, thiết lập CẤU TRÚC GIAO TIẾP gia đình liên tục (vd HỌP GIA ĐÌNH HÀNG TUẦN), và GIÁM SÁT hiệu quả.')
d.add_paragraph()

en('Efforts were made to include all caregivers in groups; we incorporated multiple engagement activities, such as motivational interviewing strategies, to increase engagement and caregiver attendance. Additional structural modifications and activities were provided to increase attendance, such as offering caregiver sessions at flexible times (e.g., after school and in the evening). We also offered transportation / taxi vouchers, child friendly activities for siblings during the sessions, and food to increase parental attendance.')
vn('Nhằm bao gồm TẤT CẢ người chăm sóc trong nhóm, chúng tôi tích hợp nhiều hoạt động ENGAGEMENT, như chiến lược PHỎNG VẤN TẠO ĐỘNG LỰC (motivational interviewing), để tăng tương tác và sự tham dự. Các điều chỉnh cấu trúc + hoạt động bổ sung được cung cấp để tăng tham dự: tổ chức buổi caregiver vào THỜI GIAN LINH HOẠT (sau giờ học, buổi tối), VOUCHER ĐI LẠI/taxi, hoạt động THÂN THIỆN VỚI TRẺ EM cho anh/chị/em trong buổi học, và THỨC ĂN.')
d.add_paragraph()

add_h('2.3. Thành phần GIÁO VIÊN (Cheer Teachers)', 2)
en('Finally, the EACP Teacher Component targeted identified "cheer teachers" for the Coping Power intervention to support their generalization of skills and to monitor progress. A teacher was identified for each participating student. These teachers engaged in three 30-min meetings during the school year with an EACP clinician.')
vn('Cuối cùng, thành phần GIÁO VIÊN của EACP nhắm vào các "CHEER TEACHERS" được xác định cho can thiệp Coping Power để hỗ trợ NGOẠI SUY KỸ NĂNG và theo dõi tiến độ. Mỗi HS tham gia có 1 giáo viên được xác định. Các giáo viên này tham gia 3 BUỔI HỌP 30 PHÚT trong năm học với một clinician của EACP.')
d.add_paragraph()

en('These meetings included a didactic presentation on EACP, reviewing its theoretical approach, the structure of student sessions, and identification of the key skills students were learning so that teachers could support and optimize their skill use in the classroom. Teachers received instruction on how to support students in setting goals, reviewing weekly tip sheets of topics and activities, discussing goal progress and challenges, encouraging students to apply skills developed in the program, building positive relationship with students, and communicating with the school counselor and EACP clinician about student concerns and progress.')
vn('Các buổi họp này gồm một BÀI THUYẾT TRÌNH GIÁO HUẤN về EACP, ôn lại cách tiếp cận lý thuyết, cấu trúc buổi HS, và xác định các kỹ năng KEY HS đang học để giáo viên có thể hỗ trợ và tối ưu việc sử dụng kỹ năng trong lớp. Giáo viên nhận hướng dẫn cách hỗ trợ HS đặt mục tiêu, xem lại các BẢNG MẸO HÀNG TUẦN về chủ đề và hoạt động, thảo luận tiến độ và thách thức, KHUYẾN KHÍCH HS áp dụng kỹ năng, xây dựng QUAN HỆ TÍCH CỰC với HS, và GIAO TIẾP với cố vấn học đường và clinician EACP về các vấn đề và tiến độ của HS.')
d.add_paragraph()

# ============ METHODS ============
page(4)
add_h('3. PHƯƠNG PHÁP (Methods)', 1)

add_h('3.1. Thiết kế nghiên cứu', 2)
en('We conducted a school-level group RCT (Murray, 1998) to evaluate the effects of EACP on 7th grade students\' behavior in 40 middle schools in Alabama and Maryland, with 20 schools enrolled in each state. The total sample of 40 schools was recruited into the project through district partnerships with the researchers.')
vn('Chúng tôi tiến hành RCT NHÓM CẤP TRƯỜNG (school-level group RCT) (Murray, 1998) để đánh giá hiệu ứng của EACP lên hành vi HS lớp 7 tại 40 trường THCS ở ALABAMA và MARYLAND, mỗi bang 20 trường. Tổng mẫu 40 trường được tuyển vào dự án qua hợp tác với các quận với nhà nghiên cứu.')
d.add_paragraph()

en('Half of the schools in each state were randomized to implement the intervention and half were assigned to school as usual, just once, prior to the screening of the first year of the project. Schools were randomized to either a control or intervention condition to avoid contamination effects that could occur if students were randomized within schools. We created 20 matched pairs of schools and randomly assigned schools within pairs to EACP and control.')
vn('MỘT NỬA các trường ở mỗi bang được PHÂN NGẪU NHIÊN triển khai can thiệp và một nửa được ấn định cho "trường-bình-thường" (school as usual), CHỈ MỘT LẦN, trước khi sàng lọc năm thứ nhất của dự án. Các trường được phân ngẫu nhiên vào điều kiện CHỨNG hoặc CAN THIỆP để tránh hiệu ứng nhiễm chéo (contamination effects) có thể xảy ra nếu HS được phân ngẫu nhiên trong cùng trường. Chúng tôi tạo 20 CẶP TRƯỜNG TƯƠNG ĐỒNG và phân ngẫu nhiên trong từng cặp.')
d.add_paragraph()

en('Participating schools had a mean enrollment of 569.03 students (SD = 238.11) and racially diverse student populations (e.g., M = 58.02% African American, SD = 34.44). In addition, schools reported, on average, that 68.06% (SD = 26.33) of students in their school were eligible for free or reduced-price meals.')
vn('Các trường tham gia có sĩ số TB 569,03 HS (SD = 238,11) và quần thể HS đa dạng SẮC TỘC (vd TB 58,02 % người Mỹ gốc Phi, SD = 34,44). Ngoài ra, các trường báo cáo TB 68,06 % (SD = 26,33) HS đủ điều kiện ĂN TRƯA MIỄN PHÍ HOẶC GIẢM GIÁ — chỉ báo SES THẤP.')
d.add_paragraph()

en('Three consecutive cohorts of approximately 120 students each year were included in the intervention condition and contrasted to the 120 control students each year. Data were collected from intervention and control students at three time points over the course of two years (one intervention year and one follow-up year). The students identified for inclusion in the study were selected from the approximately 25% of rising seventh graders in the sample who evidenced the highest levels of aggression relative to their peers, according to teacher ratings.')
vn('BA NHÓM (cohort) liên tiếp khoảng 120 HS mỗi năm được đưa vào điều kiện can thiệp và so với 120 HS đối chứng mỗi năm. Dữ liệu thu thập tại 3 THỜI ĐIỂM trong 2 NĂM (1 năm can thiệp + 1 năm theo dõi). HS được đưa vào nghiên cứu được chọn từ khoảng 25 % HS sắp lên lớp 7 trong mẫu — những HS có MỨC ĐỘ GÂY HẤN cao nhất so với bạn bè, theo đánh giá của giáo viên.')
d.add_paragraph()

# ============ RESULTS ============
page(10)
add_h('4. KẾT QUẢ (Results) — các phát hiện chính', 1)

add_h('4.1. Kết quả theo báo cáo của giáo viên (BASC-2 teacher-rated)', 2)
en('For teacher-reported outcomes, intent to treat (ITT) results indicated that EACP was associated with a decrease in externalizing problems over time. Exploration of effect modification suggested that girls in the EACP condition demonstrated statistically significant decreases in school problems over time, as well as a baseline by intervention effect whereby students and schools with greater baseline difficulties demonstrated decreased internalizing problems over time.')
vn('Đối với KẾT QUẢ giáo viên báo cáo, phân tích ITT cho thấy EACP có liên quan đến GIẢM các vấn đề NGOẠI HOÁ theo thời gian. Khám phá hiệu ứng điều biến: NỮ trong EACP có giảm có ý nghĩa thống kê các vấn đề HỌC ĐƯỜNG (school problems); cùng với hiệu ứng baseline-by-intervention: HS và trường có khó khăn baseline cao GIẢM được vấn đề NỘI HOÁ (internalizing) theo thời gian.')
d.add_paragraph()

add_h('4.2. Kết quả tự báo cáo của HS (BASC-2 student-rated)', 2)
en('Results indicated that EACP was associated with a small, but statistically significant increase in Internalizing Problems over time (γ_EACP = 0.70, p = .045, d_GMA-RAW = 0.13). There were no other significant ITT effects of EACP on growth in student-reported BASC-2 outcomes.')
vn('Kết quả cho thấy EACP có liên quan đến TĂNG NHỎ nhưng có ý nghĩa thống kê các vấn đề NỘI HOÁ theo thời gian (γ_EACP = 0,70; p = 0,045; d_GMA-RAW = 0,13). KHÔNG có các hiệu ứng ITT đáng kể khác của EACP lên các outcome BASC-2 tự báo cáo.')
d.add_paragraph()

en('Additional analyses examining effect modification resulted in one statistically significant cross-level interaction between intervention status and sex (γ_EACP*Sex = 1.83, p = .048), which indicated that girls in EACP demonstrated more favorable self-reported Personal Adjustment outcomes.')
vn('Phân tích bổ sung kiểm tra hiệu ứng điều biến cho thấy MỘT TƯƠNG TÁC chéo cấp có ý nghĩa giữa tình trạng can thiệp và giới (γ_EACP*Sex = 1,83; p = 0,048), chỉ ra NỮ trong EACP có outcome TỰ ĐIỀU CHỈNH CÁ NHÂN (Personal Adjustment) tự báo cáo TỐT HƠN.')
d.add_paragraph()

# ============ DISCUSSION ============
page(12)
add_h('5. THẢO LUẬN (Discussion) — Các điểm cốt lõi', 1)

en('This study aimed to contribute to the knowledge base on Coping Power, as there has been limited examination of the impacts of this and other prevention programs during the middle school years. These results suggest that the structured social-cognitive indicated prevention program, which includes youth, caregiver, and teacher elements, can be adapted for use with early adolescents, and have a significant main effect on teacher ratings of externalizing problems through a 1-year follow up.')
vn('Nghiên cứu này nhằm đóng góp vào BASE KIẾN THỨC về Coping Power, vì còn ít kiểm tra về tác động của chương trình này và các chương trình phòng ngừa khác TRONG GIAI ĐOẠN THCS. Kết quả gợi ý rằng chương trình phòng ngừa "INDICATED" có cấu trúc xã hội-nhận thức, gồm các thành phần HS, người chăm sóc, và giáo viên, có thể THÍCH ỨNG để dùng với VTN sớm, và có hiệu ứng chính có ý nghĩa lên đánh giá của giáo viên về vấn đề NGOẠI HOÁ tới 1 năm theo dõi.')
d.add_paragraph()

en('Notably, 8th grade teachers, who completed the follow-up assessment were unaware of the intervention condition. There was also an interaction between sex and intervention on School Problems over time, such that girls in EACP demonstrated decreases in School Problems over time.')
vn('Đáng chú ý, các GIÁO VIÊN LỚP 8 (thực hiện đánh giá follow-up) KHÔNG BIẾT về điều kiện can thiệp (blinded). Cũng có tương tác giữa GIỚI và CAN THIỆP trên School Problems theo thời gian: NỮ trong EACP có giảm các vấn đề học đường theo thời gian.')
d.add_paragraph()

en('In this context, then, it is encouraging that EACP can help reduce the escalation of externalizing behaviors among at-risk students during this developmental period; moreover, these effects are apparent for boys and girls, and for Black and White students. This developmentally adapted version of Coping Power likely reduces the students\' externalizing problems by addressing the developmentally salient emotional regulation, decision-making, and peer relationship skills and the caretakers\' monitoring and consistent discipline skills.')
vn('Trong bối cảnh này, đáng MỪNG là EACP có thể giảm leo thang các hành vi ngoại hoá ở HS có nguy cơ trong giai đoạn phát triển này; hơn nữa, các hiệu ứng này thấy được ở CẢ NAM VÀ NỮ, và ở HS DA ĐEN VÀ DA TRẮNG. Phiên bản EACP thích ứng phát triển của Coping Power có thể giảm vấn đề ngoại hoá bằng cách đề cập đến các KỸ NĂNG quan trọng về phát triển: ĐIỀU HOÀ CẢM XÚC, RA QUYẾT ĐỊNH, và QUAN HỆ BẠN BÈ; cùng với KỸ NĂNG GIÁM SÁT và KỶ LUẬT NHẤT QUÁN của người chăm sóc.')
d.add_paragraph()

en('Some children with externalizing problems have low levels of subjective anxiety, are less concerned about consequences of their behavior and thus have poorer outcomes. Modest levels of anxiety and depression have been associated with better intervention response among children with disruptive behaviors, and, in the present study, modest increases in Internalizing Problems may indicate children\'s increased awareness of consequences of externalizing behaviors.')
vn('Một số trẻ có vấn đề ngoại hoá có MỨC ĐỘ LO ÂU CHỦ QUAN THẤP, ÍT lo lắng về hậu quả hành vi và do đó có kết quả XẤU HƠN. Mức độ lo âu và trầm cảm KHIÊM TỐN đã được liên kết với phản ứng can thiệp TỐT HƠN ở trẻ có hành vi gây rối, và trong nghiên cứu này, các tăng KHIÊM TỐN trong vấn đề nội hoá có thể chỉ ra trẻ TĂNG NHẬN THỨC về hậu quả của hành vi ngoại hoá.')
d.add_paragraph()

# ============ PHẢN BIỆN ============
add_h('PHẢN BIỆN CHI TIẾT (em viết — không có trong bài gốc)', 1, RED)
crit_points = [
    ('① RCT cluster 40 trường — bằng chứng cấp cao',
     'Đa trung tâm 2 bang (Alabama + Maryland), 20 cặp trường tương đồng, randomize ở cấp trường — methodology chuẩn vàng (Murray 1998). NIMH funding, peer-review J School Psychology. Có 1-year follow-up với 8th grade teachers blinded → giảm bias. Dùng BASC-2 — thang đo lâm sàng chuẩn.'),
    ('② Đối tượng GỐC là HS GÂY HẤN — không phải HS lo âu thuần',
     'Sàng lọc 25 % HS có MỨC GÂY HẤN cao nhất qua giáo viên. Outcome chính là EXTERNALIZING (hành vi gây hấn), không phải lo âu/trầm cảm. Cảnh báo cho áp dụng VN: nếu thầy muốn giảm LO ÂU thì EACP không phải first-line — cần CBT cho lo âu (CAMS, Walder, Matsumoto).'),
    ('③ Hiệu ứng đặc biệt với NỮ — phát hiện mới',
     'NỮ trong EACP cải thiện School Problems + Personal Adjustment có ý nghĩa thống kê. NAM không có hiệu ứng tương đương. Có thể do: (a) cấu trúc 25 buổi nhóm phù hợp NỮ hơn (verbal); (b) nội dung tình bạn/tình cảm nhạy với NỮ. Cho VN: EACP có thể MẠNH cho HS NỮ hơn HS NAM.'),
    ('④ Tăng nội hoá tự báo cáo — paradox',
     'γ = 0,70 (tăng), p = 0,045, d = 0,13 (rất nhỏ nhưng significant). Tác giả lý giải: HS sau EACP NHẬN THỨC HƠN về hậu quả → tự báo cáo lo âu tăng. Đây có thể là "AWARENESS effect" tích cực, không phải tác hại. Nhưng cần verify trong các RCT khác.'),
    ('⑤ Bias quần thể — 69,8 % African American + SES thấp',
     'Mẫu chủ yếu HS Mỹ gốc Phi (Alabama + Maryland — 2 bang Bắc-Nam tương đối nghèo). 68 % HS đủ ăn miễn phí. Áp dụng cho VN cần adapt: (a) khái niệm cyberbullying + deviant peer khác; (b) Family meeting structures khó với gia đình VN; (c) cheer teacher concept cần adapt.'),
    ('⑥ EACP-VN: gợi ý adapt',
     'Nếu adapt cho HS THCS VN: (1) thay "anger" thành "lo âu/stress" cho phù hợp đề tài; (2) thay PICC bằng GROW model hoặc tương đương VN; (3) bỏ "cyberbullying" hoặc thay bằng "tin nhắn/mạng xã hội"; (4) thay "cheer teacher" thành "giáo viên chủ nhiệm"; (5) cho cha mẹ — bỏ "family meeting" thay bằng "trò chuyện hằng tuần".'),
    ('⑦ Đối chiếu với DB của em',
     'EACP là can thiệp PHÒNG NGỪA (indicated prevention) — khác Walder QT040 DMHI cho SAD đã có triệu chứng. Cũng khác Maya app QT043 (CBT cho người trẻ 18-25). EACP unique: school-based + group + 3 thành phần (HS+cha mẹ+giáo viên) — phù hợp triển khai trường VN nhất.'),
]
for ttl, body in crit_points:
    p = d.add_paragraph(); rr = p.add_run(ttl); rr.bold = True; rr.font.color.rgb = RED
    p2 = d.add_paragraph(); rr2 = p2.add_run(body); rr2.font.color.rgb = RED

d.add_paragraph()

foot = d.add_paragraph()
fr = foot.add_run('Bản dịch song ngữ Anh-Việt — biên soạn 29/04/2026. Phần References giữ nguyên tiếng Anh trong PDF gốc (theo Nguyên tắc dịch v2 — Nguyên tắc 2). Số liệu: γ_EACP = 0,70, d = 0,13 verified từ p10 PDF gốc.')
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
