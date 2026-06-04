# -*- coding: utf-8 -*-
"""Chinh dinh dang LA v2 — CONSERVATIVE: bao toan alignment goc.
- Margins section: 2.5/2.5/3.5/2.0 cm
- Font run: Times New Roman (chi sua font name)
- Font size: 13pt cho body, 14pt H2, 16pt H1 — CHI khi co style Heading X tu goc
- Line spacing: 1.5 cho body, neu chua duoc set
- KHONG sua alignment, KHONG sua indent (giu nguyen quyet dinh goc)
- TOC field
- Xoa watermark + metadata
28/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
# Re-do from v3_6 (clean state)
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_6_FixAnderson_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

print(f"SRC: {SRC}")
print(f"OUT: {OUT}")
print()

d = Document(SRC)

# ============================================================
# STEP 1: Section margins + watermarks
# ============================================================
print("--- Step 1: Margins + watermark removal ---")
removed_wm = 0
for sec in d.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2.0)
    for hf in [sec.header, sec.first_page_header, sec.even_page_header,
               sec.footer, sec.first_page_footer, sec.even_page_footer]:
        try:
            for elem in list(hf._element.iter()):
                if elem.tag in (qn('w:pict'), qn('w:object'), qn('w:drawing')):
                    if elem.getparent() is not None:
                        elem.getparent().remove(elem)
                        removed_wm += 1
        except: pass
print(f"  Margins 2.5/2.5/3.5/2.0 set on {len(d.sections)} sections")
print(f"  Removed {removed_wm} watermark elements from headers/footers")
print()

# ============================================================
# STEP 2: CONSERVATIVE paragraph formatting
# Only change: font name (TNR), font size if heading, line spacing if body
# DO NOT change: alignment, indent, bold, italic, color
# ============================================================
print("--- Step 2: Conservative paragraph format ---")

import re
def get_heading_level_from_style(p):
    """Only use Word style to identify heading level. Return 1-4 or 0."""
    style_name = p.style.name if p.style else ''
    if 'Heading 1' in style_name or style_name == 'Title':
        return 1
    if 'Heading 2' in style_name:
        return 2
    if 'Heading 3' in style_name:
        return 3
    if 'Heading 4' in style_name or 'Heading 5' in style_name:
        return 4
    return 0


def is_chuong_heading(text):
    """Detect CHƯƠNG <num> short heading. Must be all-caps, short, structural."""
    t = text.strip()
    if len(t) > 100:
        return False
    if re.match(r'^CHƯƠNG\s+[\dIVX]+', t):
        return True
    upper = t.upper()
    h1_keywords = ['MỞ ĐẦU', 'KẾT LUẬN', 'KIẾN NGHỊ', 'TÀI LIỆU THAM KHẢO',
                   'PHỤ LỤC', 'MỤC LỤC', 'LỜI CAM ĐOAN', 'LỜI CẢM ƠN',
                   'DANH MỤC CÔNG TRÌNH ĐÃ CÔNG BỐ',
                   'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
                   'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
                   'KẾT LUẬN VÀ KIẾN NGHỊ']
    for kw in h1_keywords:
        if upper == kw or upper.startswith(kw + ':') or upper.startswith(kw + '.'):
            return True
    return False


n_h1_style = n_h2_style = n_h3_style = n_h4_style = 0
n_chuong = 0
n_body_font = 0

for p in d.paragraphs:
    text = p.text.strip()
    if not text:
        # Skip empty paragraphs - leave alone
        continue

    level = get_heading_level_from_style(p)
    is_chuong = is_chuong_heading(text)

    if level == 1 or is_chuong:
        # CHƯƠNG / front-back matter heading: 16pt bold, KEEP alignment
        for r in p.runs:
            if r.font.name != 'Times New Roman':
                r.font.name = 'Times New Roman'
            r.font.size = Pt(16)
            r.bold = True
        if is_chuong:
            n_chuong += 1
        else:
            n_h1_style += 1
    elif level == 2:
        # H2: 14pt bold, KEEP alignment + indent
        for r in p.runs:
            if r.font.name != 'Times New Roman':
                r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.bold = True
        n_h2_style += 1
    elif level == 3:
        # H3: 13pt bold
        for r in p.runs:
            if r.font.name != 'Times New Roman':
                r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            r.bold = True
        n_h3_style += 1
    elif level == 4:
        # H4: 13pt italic bold
        for r in p.runs:
            if r.font.name != 'Times New Roman':
                r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            r.bold = True
            r.italic = True
        n_h4_style += 1
    else:
        # BODY: only standardize font name + size
        # KEEP alignment, KEEP indent, KEEP bold/italic/color
        for r in p.runs:
            if r.font.name != 'Times New Roman':
                r.font.name = 'Times New Roman'
            # Only set size if not yet set or is wrong (preserve special small/big)
            if r.font.size is None:
                r.font.size = Pt(13)
            elif r.font.size.pt < 10 or r.font.size.pt > 16:
                # Likely error — clamp to 13
                r.font.size = Pt(13)
            # else keep existing size
        n_body_font += 1
        # Line spacing: set to 1.5 ONLY if not set
        pf = p.paragraph_format
        if pf.line_spacing is None:
            pf.line_spacing = 1.5
            pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE

print(f"  CHƯƠNG/front-back headings: {n_chuong} (giu align goc, 16pt bold)")
print(f"  Heading 1 style: {n_h1_style} (16pt bold)")
print(f"  Heading 2 style: {n_h2_style} (14pt bold)")
print(f"  Heading 3 style: {n_h3_style} (13pt bold)")
print(f"  Heading 4 style: {n_h4_style} (13pt italic bold)")
print(f"  Body paragraphs (chi sua font): {n_body_font}")
print()

# ============================================================
# STEP 3: Table cells — only font name
# ============================================================
print("--- Step 3: Table cells (only font name) ---")
n_cells = 0
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for r in p.runs:
                    if r.font.name != 'Times New Roman':
                        r.font.name = 'Times New Roman'
                n_cells += 1
print(f"  {n_cells} cells fixed font (keep size + alignment)")
print()

# ============================================================
# STEP 4: TOC field
# ============================================================
print("--- Step 4: TOC field ---")

def make_toc_paragraph(parent_para):
    r = parent_para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r._r.append(instrText)
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fldChar_sep)
    instrResult = OxmlElement('w:t')
    instrResult.text = 'Nhấn F9 hoặc chuột phải > Update Field để cập nhật Mục lục.'
    r._r.append(instrResult)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar_end)


toc_para_idx = None
for i, p in enumerate(d.paragraphs[:200]):
    if p.text.strip().upper() in ('MỤC LỤC', 'MUC LUC'):
        toc_para_idx = i
        break

if toc_para_idx is not None:
    print(f"  Tim thay MỤC LỤC tai para {toc_para_idx}. Chen TOC ngay sau.")
    toc_heading_p = d.paragraphs[toc_para_idx]
    # Don't touch alignment of MỤC LỤC heading - it's already center
    new_p_xml = OxmlElement('w:p')
    toc_heading_p._element.addnext(new_p_xml)
    from docx.text.paragraph import Paragraph
    toc_p_obj = Paragraph(new_p_xml, toc_heading_p._parent)
    make_toc_paragraph(toc_p_obj)
else:
    print("  Khong tim thay MỤC LỤC. Khong chen TOC (de tranh dao loan).")
print()

# ============================================================
# STEP 5: Metadata
# ============================================================
print("--- Step 5: Clean metadata ---")
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
print("  Cleared")
print()

# ============================================================
# Save
# ============================================================
d.save(OUT)
print(f"--- Done ---")
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)//1024} KB")
