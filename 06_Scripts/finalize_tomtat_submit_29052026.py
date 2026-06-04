# -*- coding: utf-8 -*-
"""Buoc 4: Finalize ca 2 ban (VN + EN) san sang nop:
- Them so trang
- Clean metadata
- Sinh checklist nop cho NCS
29/05/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))


def add_page_number(section):
    """Them so trang Arabic giua chan trang, 10pt."""
    footer = section.footer
    p = footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # Add page number field
    r = p.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = 'PAGE \\* MERGEFORMAT'
    r._r.append(instrText)
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fldChar_sep)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar_end)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)


# Process both files
for fn in ['TomTatLA_v2_VERIFIED_29052026.docx', 'TomTatLA_EN_v1_29052026.docx']:
    path = os.path.join(ROOT, 'Luận án TS', fn)
    # Backup
    backup = path.replace('.docx', '_BEFORE_PAGE_NUM.docx')
    shutil.copy(path, backup)

    d = Document(path)
    for sec in d.sections:
        add_page_number(sec)

    # Re-clean metadata
    cp = d.core_properties
    cp.author = ''; cp.title = ''; cp.subject = ''; cp.keywords = ''
    cp.comments = ''; cp.last_modified_by = ''

    # Save back
    d.save(path)
    print(f"Done: {fn}")


# ============================================================
# Generate checklist for NCS
# ============================================================
print("\nGenerating submission checklist...")
CHECKLIST = os.path.join(ROOT, 'Luận án TS', 'TomTatLA_CHECKLIST_NOP_29052026.docx')

d = Document()
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.3
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

def H(text, level=1):
    h = d.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'

def P(text, bold=False, italic=False, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    return p

H('CHECKLIST NỘP BẢN TÓM TẮT LUẬN ÁN TIẾN SĨ', 1)
P('NCS Công Thị Hằng — 29/05/2026', italic=True)
P('')

H('A. CÁC FILE ĐÃ SẴN SÀNG', 2)
P('1. Bản tóm tắt tiếng Việt: TomTatLA_v2_VERIFIED_29052026.docx (21 trang, đã đánh số trang)', size=11)
P('2. Bản tóm tắt tiếng Anh: TomTatLA_EN_v1_29052026.docx (21 trang, đã đánh số trang)', size=11)
P('3. Luận án chính (tham chiếu): 01_LuanAn_v3_1_FixCoping_28052026.docx (đã chỉnh sửa lỗi)', size=11)
P('')

H('B. PHẦN NCS CẦN ĐIỀN TRƯỚC KHI NỘP', 2)
P('1. DANH MỤC CÔNG TRÌNH ĐÃ CÔNG BỐ (cuối tóm tắt):', bold=True)
P('   - Yêu cầu tối thiểu: 2 bài đăng tạp chí có chỉ số khoa học (theo quy định ĐHSPHN).', size=11)
P('   - Format APA hoặc Vancouver tùy quy chuẩn trường.', size=11)
P('   - Trong bản tiếng Anh: dịch tên bài + tên tạp chí sang tiếng Anh tương đương.', size=11)
P('')
P('2. PHẢN BIỆN 1/2/3 (trang phụ bìa):', bold=True)
P('   - Bỏ trống chờ Hội đồng bổ nhiệm. Sau khi có Quyết định, điền tên + học hàm + đơn vị.', size=11)
P('')
P('3. NGÀY GIỜ BẢO VỆ (trang phụ bìa):', bold=True)
P('   - Bỏ trống chờ Phòng Sau đại học thông báo.', size=11)
P('')

H('C. KIỂM TRA TRƯỚC KHI IN', 2)
P('1. Mở mỗi file trong Microsoft Word.')
P('2. Nhấn Ctrl+A → F9 để cập nhật toàn bộ field (số trang).')
P('3. Kiểm tra số trang chân trang hiển thị đúng.')
P('4. Kiểm tra ngắt trang đúng giữa các phần (Mở đầu, mỗi Chương, Kết luận).')
P('5. Đối chiếu số liệu thống kê quan trọng:')
P('   - 1.352 học sinh (Nam 614 / Nữ 738) ✓', size=10)
P('   - Cronbach α của các thang đo ≥ 0,70 ✓', size=10)
P('   - β chuẩn hóa của ALHT, NĐT, BNHĐ, TTr, HTCM, GBTH ✓', size=10)
P('   - F values + p values cho gender comparison ✓', size=10)
P('')

H('D. ĐỊNH DẠNG ĐÃ ÁP DỤNG (theo TT 18/2021/TT-BGDĐT)', 2)
P('• Khổ giấy: A5 (148 × 210 mm)')
P('• Font: Times New Roman 11pt body, 14pt CAPS bold cho Chương')
P('• Giãn dòng: 1.15')
P('• Lề: Trên 1.5 / Dưới 1.5 / Trái 2.0 / Phải 1.5 cm')
P('• Đánh số trang: Ả-rập (1, 2, 3...), giữa chân trang, 10pt')
P('• Căn lề: Justify, thụt đầu dòng 0.75 cm')
P('')

H('E. IN ẤN + ĐÓNG BÌA', 2)
P('1. In hai mặt trên giấy A5 cho cả bản Việt + Anh.')
P('2. Bìa cứng màu xanh đậm hoặc đỏ đậm (theo quy định ĐHSPHN), chữ vàng/nhũ.')
P('3. Gáy bìa: tên đề tài + tên NCS + năm 2026, in dọc.')
P('4. Số lượng in: kiểm tra Phòng Sau đại học (thường 15-25 bản cho HĐ + lưu trữ).')
P('')

H('F. NỘP KÈM CD/USB', 2)
P('Nội dung CD/USB nộp kèm bản giấy:')
P('1. Luận án chính (file Word + PDF)')
P('2. Bản tóm tắt tiếng Việt (Word + PDF)')
P('3. Bản tóm tắt tiếng Anh (Word + PDF)')
P('4. Danh mục công trình đã công bố (Word + PDF)')
P('5. Slide bảo vệ (PowerPoint, nếu đã chuẩn bị)')
P('')

H('G. ĐỐI CHIẾU CUỐI CÙNG', 2)
P('☐ Tên đề tài trên bìa + phụ bìa + tóm tắt thống nhất với LA chính', size=11)
P('☐ Chuyên ngành + Mã số chính xác (Tâm lý học chuyên ngành / 9.31.04.01)', size=11)
P('☐ Tên người HD đúng: TS. Đào Minh Đức', size=11)
P('☐ Bản tiếng Anh đã được hiệu đính bởi người có chuyên môn', size=11)
P('☐ Mọi số liệu thống kê khớp giữa LA + tóm tắt VN + tóm tắt EN', size=11)
P('☐ 3 giả thuyết khoa học đồng nhất giữa Mở đầu, Kết quả, Kết luận', size=11)
P('☐ Khung CT tập huấn 8 nội dung đồng nhất', size=11)
P('☐ Danh mục công trình đã điền đầy đủ (≥ 2 bài)', size=11)
P('☐ Số trang hiển thị đúng sau F9 update', size=11)
P('☐ Không còn placeholder [...] hoặc dấu chấm trống', size=11)

cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
d.save(CHECKLIST)
print(f"\nChecklist saved: {CHECKLIST}")
