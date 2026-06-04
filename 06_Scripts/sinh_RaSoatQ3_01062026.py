# -*- coding: utf-8 -*-
"""Sinh bao cao ra soat bai Q3 (PLOS ONE descriptive paper).
Logic flow + every claim + every data point.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'RaSoat_Q3_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.3


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=False):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def V(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('✓ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)

def X(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('⚠ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


# ============================================================
H1('BÁO CÁO RÀ SOÁT BÀI Q3 — LOGIC FLOW + CLAIMS')
P('Rà soát chi tiết: luồng dẫn dắt + mọi câu khẳng định + mọi dữ liệu',
  italic=True)
P('Ngày: 01/06/2026 — Phương án A (PLOS ONE descriptive paper)', italic=True)


# ============================================================
H1('PHẦN 1 — LOGIC FLOW REVIEW (luồng dẫn dắt Q3)')

H2('1.1 Tổng quan dòng chảy Introduction')

H3('Hiện tại (outline v2)')
P('§1 → Adolescent anxiety burden + Asian context')
P('§2 → Vietnam gap: limited item-level + normative data')
P('§3 → Present study + 3 RQs')

H3('Vấn đề phát hiện')
X('Issue Q3-1: §1 + §2 quá ngắn — Q3 cần introduce: WHY item-level analysis '
  'matters (cite Chorpita 2000 original RCADS, why subset adaptation, why '
  'normative data for VN clinical use)')
X('Issue Q3-2: Research Questions (RQ1-3) chưa testable — chỉ description '
  'chung. Cần phát biểu cụ thể.')
X('Issue Q3-3: Missing RATIONALE for descriptive approach — PLOS ONE reviewers '
  'thường hỏi "Why descriptive instead of inferential?". Cần justify: '
  '(1) normative data gap; (2) screening tool development for VN.')
X('Issue Q3-4: Citation Anderson 2025, V-NAMHS 2022 hiện chỉ mention — '
  'cần CONTEXT VIETNAMESE specific (V-NAMHS reports 2,3% với DISC-5 — '
  'chênh lệch lớn với các study DASS-21 → KHOẢNG TRỐNG cho RCADS adaptation)')

H3('Đề xuất sửa Introduction Q3')
V('§1 (expand): Adolescent anxiety burden (Xu 2021 N=373,216; Anderson 2025 '
  'narrative; GBD ASEAN 2025) + WHY DSM-5 subtypes matter (different treatments) '
  '→ ~250 từ')
V('§2 (expand): Vietnam under-researched item-level + normative data gap; '
  'V-NAMHS 2022 (2,3% with DISC-5) vs Hoang Trung Hoc 2025 (higher with DASS); '
  'methodological need for RCADS Vietnamese adaptation → ~300 từ')
V('§3 (new): RATIONALE for descriptive: (a) screening tool development '
  '(b) normative comparison data for VN clinicians (c) grade trajectory '
  'context → ~150 từ')
V('§4 (revised): 3 testable RQs:', )
V('   RQ1: "What is the item-level distribution of GAD/SAD/SocAD symptoms '
  'in Vietnamese lower secondary students?"',)
V('   RQ2: "Do anxiety symptoms differ by gender and grade level?"',)
V('   RQ3: "Which specific RCADS items best identify high-anxiety students '
  'for screening purposes?"',)


H2('1.2 Logic flow Methods → Results → Discussion')

H3('Vấn đề phát hiện')
X('Issue Q3-5: Methods 2.4 chỉ list techniques (descriptive, ANOVA) mà KHÔNG '
  'specify multiple comparison correction (Bonferroni? Holm? FDR?). Cần.')
X('Issue Q3-6: Methods 2.5 (Ethics) chỉ ghi "identical to Q1" — Q3 standalone '
  'paper cần full ethics statement, không thể "see Q1".')
X('Issue Q3-7: Results 3.1-3.6 đầy đủ items nhưng THIẾU summary table - '
  'PLOS ONE thích visualizations + summary stats.')
X('Issue Q3-8: Discussion 4.3 "brief mention" gender invariance — PLOS ONE '
  'reviewers có thể yêu cầu more detail. Cần balance giữa "reserve for Q1" '
  '(avoid self-plagiarism) và "answer reviewer questions".')
X('Issue Q3-9: Cross-reference Q3↔Q1 strategy chưa rõ. Khi Q1 submit + accept '
  'trước, Q3 cite ra sao? Hoặc submit cùng lúc (recommend explanation)?')

H3('Đề xuất sửa Methods Q3')
V('2.4 Analytic strategy: add "Bonferroni correction for multiple comparisons '
  '(α = 0,05/k where k = number of tests); effect size: Cohen d (gender), '
  'η² (grade)"')
V('2.5 Ethics: write FULL statement (don\'t reference Q1) — "Approved by HNUE '
  'IRB [number/date]; parental consent + student assent; data anonymized"')

H3('Đề xuất sửa Results + Discussion Q3')
V('3 add: "Table 6: Summary of key findings + screening cut-points" — '
  'PLOS ONE-friendly visual')
V('Figure 1: Already proposed — grade trajectory plot for 3 subtypes (key visual)')
V('4.3 (gender): Brief 2-3 sentences với cite Q1 paper if đã submit (or '
  '"Future research will explore gender invariance in greater detail [Cong & Dao, '
  '2026 in prep]")')
V('Cross-reference strategy: Submit Q1 first → wait acceptance → submit Q3 with '
  'citation to Q1. OR Submit cùng lúc with explicit note in Methods: '
  '"Companion paper [Cong et al., under review] addresses the integrated '
  'mechanism analysis."')


d.add_page_break()


# ============================================================
H1('PHẦN 2 — CLAIMS VERIFICATION Q3')

H2('2.1 Introduction citations Q3')
V('Xu 2021 → VERIFIED PDF QT010')
V('Wen 2020 → VERIFIED PDF QT008')
V('Saikia 2023 → VERIFIED PDF 11_Saikia_2023_IJCM.pdf (real Saikia)')
V('Bhardwaj 2020 → VERIFIED PDF QT011')
V('Anderson 2025 → VERIFIED PDF QT014')
V('V-NAMHS 2022 → VERIFIED PDF V-NAMHS_2022.pdf')
V('Hoang Trung Hoc 2025 → VERIFIED PDF VN014_HoangTrungHoc_2025_VN_COVID.pdf')
V('Nakie 2022 → VERIFIED PDF QT006')

H2('2.2 Methods claims Q3')
V('1.352 HS → VERIFIED (same dataset)')
V('RCADS Chorpita 2000 → VERIFIED LA P959 (15 items: 7+4+4)')
V('4-point Likert → VERIFIED LA P959')
V('Cronbach α + McDonald ω → VERIFIED LA P962+')

H2('2.3 Results data Q3 - ALL ITEMS từ Thực trạng (verified)')

H3('GAD (Bảng 1)')
V('RCADS4 M=64,28 SD=29,69 rank 1 → VERIFIED Bảng 1')
V('RCADS13 M=59,62 SD=35,86 rank 2 → VERIFIED Bảng 1')
V('RCADS8 M=59,02 SD=33,48 rank 3 → VERIFIED Bảng 1')
V('RCADS30 M=57,69 SD=31,26 rank 4 → VERIFIED Bảng 1')
V('RCADS12 M=55,13 SD=33,84 rank 5 → VERIFIED Bảng 1')
V('RCADS1 M=49,14 SD=30,72 rank 6 → VERIFIED Bảng 1')
V('RCADS35 M=45,86 SD=33,83 rank 7 → VERIFIED Bảng 1')
V('Subtype mean M=55,82 → VERIFIED outline + Thực trạng')

H3('SAD (Bảng 2)')
V('RCADS46 M=27,88 rank 1 → VERIFIED Bảng 2')
V('RCADS17 M=25,49 rank 2 → VERIFIED Bảng 2')
V('RCADS5 M=25,35 rank 3 → VERIFIED Bảng 2')
V('RCADS45 M=21,52 rank 4 → VERIFIED Bảng 2')
V('Subtype mean M=25,06 → VERIFIED')

H3('SocAD (Bảng 3)')
V('RCADS32 M=56,98 rank 1 → VERIFIED Bảng 3')
V('RCADS43 M=49,26 rank 2 → VERIFIED Bảng 3')
V('RCADS38 M=45,32 rank 3 → VERIFIED Bảng 3')
V('RCADS20 M=42,09 rank 4 → VERIFIED Bảng 3')
V('Subtype mean M=48,41 → VERIFIED')

H3('Demographic comparisons (Bảng 4)')
V('Gender F values: 44,484; 45,984; 29,642; 0,246 → VERIFIED all')
V('p values: p<0,001 (3); p=0,620 (separation) → VERIFIED')
V('Grade trajectory SAD: 32,13 → 27,14 → 20,88 → 19,46 → VERIFIED Bảng 4')
V('Grade trajectory GAD: 54,32 → 53,65 → 55,63 → 59,79 → VERIFIED Bảng 4')
V('Grade SocAD F=4,879; p=0,002 → VERIFIED')

H2('2.4 Discussion claims Q3')
V('Anxiety academic worry dominant — derived from data ✓')
V('SocAD social judgment dominant — derived from data ✓')
V('Separation declines + GAD increases developmental pattern — VERIFIED data')

X('Discussion 4.3: Gender brief mention — chưa có cite để justify "reserve for '
  'Q1". Cần add inline note: "(see companion analysis, Cong et al., in prep)"')


d.add_page_break()


# ============================================================
H1('PHẦN 3 — TÓM TẮT 9 ISSUES Q3 + RECOMMENDATIONS')

t = d.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = '#'; hdr[1].text = 'Issue'; hdr[2].text = 'Recommendation'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

issues = [
    ('Q3-1', 'Introduction §1+§2 quá ngắn',
     'Expand với DSM-5 subtypes rationale + Vietnam normative gap'),
    ('Q3-2', 'RQ1-3 không testable',
     'Spell out 3 RQs as specific questions answerable from data'),
    ('Q3-3', 'Missing rationale for DESCRIPTIVE approach',
     'Add 1 paragraph justifying: screening + normative data + grade context'),
    ('Q3-4', 'V-NAMHS context not exploited',
     'Add comparison: V-NAMHS 2,3% DISC-5 vs Hoang DASS higher → '
     'methodological gap → RCADS Vietnamese adaptation as bridge'),
    ('Q3-5', 'Multiple comparison correction not specified',
     'Add Bonferroni or Holm correction; effect sizes (Cohen d + η²)'),
    ('Q3-6', 'Ethics statement "see Q1" inadequate',
     'Write FULL ethics statement standalone — IRB + consent + assent + '
     'anonymization'),
    ('Q3-7', 'Results missing summary visualization',
     'Add Table 6 (summary + screening cut-points) + Figure 1 (grade '
     'trajectory) — PLOS ONE-friendly'),
    ('Q3-8', 'Discussion gender invariance brief vs reviewer expectation',
     'Add 2-3 sentence justification with companion paper reference'),
    ('Q3-9', 'Cross-reference Q1↔Q3 strategy unclear',
     'Specify: (Option A) Submit Q1 first, wait accept, then Q3 cites Q1; '
     '(Option B) Submit cùng lúc với "companion paper under review" note'),
]
for issue in issues:
    row = t.add_row().cells
    for i, txt in enumerate(issue):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)


d.add_page_break()


# ============================================================
H1('PHẦN 4 — VERIFICATION SUMMARY Q3')

t2 = d.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
hdr2 = t2.rows[0].cells
hdr2[0].text = 'Category'; hdr2[1].text = 'Verified'; hdr2[2].text = 'Issues'
for c in hdr2:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

ver = [
    ('Introduction citations', '8/8 verified', '0 (all PDFs exist)'),
    ('Methods sample + scale', '4/4 verified', '0'),
    ('Results GAD item-level (7 items)', '7/7 verified', '0 (all from Thực trạng)'),
    ('Results SAD item-level (4 items)', '4/4 verified', '0'),
    ('Results SocAD item-level (4 items)', '4/4 verified', '0'),
    ('Demographic F/p values', '8/8 verified', '0'),
    ('Grade trajectory', '8 values verified', '0'),
    ('Discussion claims', '3/4 verified', '1 (companion paper ref)'),
    ('Logic flow Introduction', '—', '4 issues (expand §1+§2; testable RQs; rationale; V-NAMHS context)'),
    ('Logic flow Methods', '—', '2 issues (multiple comparison; ethics standalone)'),
    ('Logic flow Results/Discussion', '—', '3 issues (summary table; gender depth; cross-ref strategy)'),
]
for row_data in ver:
    row = t2.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if 'verified' in txt.lower() and 'issues' not in txt.lower():
                    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                elif 'issues' in txt.lower() and txt != '0':
                    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


H2('Tổng kết em đề xuất')
P('Q3 outline có ÍT issues hơn Q1 (9 vs 10) vì:', italic=True)
P('  • Descriptive approach đơn giản hơn integrated SEM', indent=False)
P('  • Tất cả data points đều verify được trực tiếp từ "Thực trạng" source', indent=False)
P('  • Không có R² integrated problem như Q1 Issue 8', indent=False)
P('  • Không có subtype-specific β table missing như Q1 Issue 9', indent=False)

P('', italic=True)
P('Critical priority Q3:', )
P('  • Q3-6 Ethics statement standalone (BLOCKING for PLOS ONE submission)', indent=False)
P('  • Q3-9 Cross-reference strategy Q1↔Q3 (CRITICAL for anti-plagiarism)', indent=False)

P('Medium priority Q3:', )
P('  • Q3-1, 2, 3, 4 (Introduction logic flow + RQs + rationale)', indent=False)
P('  • Q3-5 (Bonferroni + effect sizes)', indent=False)
P('  • Q3-7 (summary visualization)', indent=False)
P('  • Q3-8 (gender brief justification)', indent=False)

P('Em đề xuất build bố cục v3 sau khi:', italic=True)
P('  (1) Thầy + NCS chốt Q1↔Q3 submission strategy (Option A or B)', indent=False)
P('  (2) NCS xin IRB letter HNUE (cho ethics statement đầy đủ)', indent=False)
P('  (3) NCS confirm qualitative interviews count (cho Q1 Issue 6)', indent=False)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
