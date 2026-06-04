# -*- coding: utf-8 -*-
"""Sinh bao cao prioritization 19 issues Q1+Q3 + action plan.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'IssuesPriority_Q1Q3_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.8)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.25


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(text, italic=False):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic


def make_table(headers, rows, colors=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
                    if colors and i in colors and colors[i](txt):
                        r.font.color.rgb = colors[i](txt)
    return t


# ============================================================
H1('PRIORITIZATION + ACTION PLAN — 19 ISSUES Q1+Q3')
P('Phân loại 10 issues Q1 + 9 issues Q3 theo priority + ownership.', italic=True)
P('Ngày: 01/06/2026', italic=True)


# ============================================================
H2('TIER 1 — BLOCKING (4 issues): KHÔNG SUBMIT được nếu chưa fix')
P('Cần decision/data từ NCS/thầy trước khi em fix outline.', italic=True)

t1_data = [
    ('Q1-6', 'Q1 Methods 2.3', 'Qualitative interviews: số participants + transcripts + recruitment',
     'NCS', 'Confirm: n=? + sampling strategy + transcribed yet?',
     'Cannot complete Methods section'),
    ('Q1-8', 'Q1 Results 3.4', 'R²=0,284 + 0,209 là từ 2 SEPARATE models trong LA, KHÔNG phải integrated',
     'Thầy + NCS', '(A) Re-run integrated SEM với 7 predictors → 1 R²; '
     '(B) Modify claim: "Separate-model R² presented; integrated R² in companion analysis"',
     'BMC Psychiatry reviewers sẽ flag inconsistency'),
    ('Q3-6', 'Q3 Methods 2.5', 'Ethics statement "see Q1" inadequate cho standalone paper',
     'NCS', 'Xin IRB letter HNUE (retroactive nếu cần) — 2-4 tuần',
     'PLOS ONE reject submission'),
    ('Q3-9', 'Q3 anti-plagiarism', 'Cross-ref Q1↔Q3 strategy chưa quyết',
     'Thầy + NCS', '(A) Q1 first → Q3 cite Q1; '
     '(B) Submit cùng lúc với "companion paper under review" note',
     'Risk redundant publication / self-plagiarism'),
]
make_table(['ID', 'Vị trí', 'Vấn đề', 'Ownership', 'Action needed', 'Consequence nếu không fix'], t1_data)


H2('TIER 2 — HIGH PRIORITY (12 issues): Em có thể fix outline v3')
P('Cải thiện logic flow + methodological completeness. Reviewers Q1 sẽ flag '
  'nếu thiếu.', italic=True)

t2_data = [
    ('Q1-1', 'Q1 Intro §1→§2', 'Thiếu cầu nối logic burden → framework',
     'Em (outline v3)', 'Revise § với cultural specificity + integrated rationale'),
    ('Q1-2', 'Q1 Intro', 'Mixed-methods rationale absent',
     'Em', 'Add 2 sentences cuối §4 justifying qualitative'),
    ('Q1-3', 'Q1 Intro §4', 'Hypotheses không testable',
     'Em', 'Spell out 3 testable hypotheses (H1-H3 đầy đủ)'),
    ('Q1-4', 'Q1 Intro §1', 'Vietnamese cultural specificity vắng mặt',
     'Em', 'Add closing line về Confucian + family hierarchy'),
    ('Q1-5', 'Q1 Methods 2.4', 'Fit indices threshold không specify',
     'Em', 'Add: CFI ≥ 0,90; TLI ≥ 0,90; RMSEA ≤ 0,08; SRMR ≤ 0,08 '
     '(Hu & Bentler 1999); ΔCFI ≤ 0,01 (Cheung & Rensvold 2002)'),
    ('Q1-7', 'Q1 Methods', 'Mixed-methods Integration Design missing',
     'Em', 'Add: "Convergent Parallel Design (Creswell & Plano Clark 2018) + '
     'joint display matrix"'),
    ('Q1-9', 'Q1 Results 3.4', 'Subtype-specific β table missing',
     'Em (re-extract LA)', 'Build table 7 predictors × 3 outcomes (21 β) '
     'từ LA Bảng 27-42'),
    ('Q3-1', 'Q3 Intro §1+§2', 'Quá ngắn',
     'Em', 'Expand với DSM-5 subtypes rationale + Vietnam normative gap'),
    ('Q3-2', 'Q3 Intro §3', 'RQ1-3 không testable',
     'Em', 'Spell out 3 specific RQs'),
    ('Q3-3', 'Q3 Intro', 'Missing rationale descriptive approach',
     'Em', 'Add paragraph: screening + normative + grade trajectory'),
    ('Q3-4', 'Q3 Intro §2', 'V-NAMHS context chưa khai thác',
     'Em', 'Add comparison V-NAMHS 2,3% vs DASS → methodological gap'),
    ('Q3-5', 'Q3 Methods 2.4', 'Bonferroni + effect sizes missing',
     'Em', 'Add: Bonferroni correction + Cohen d (gender) + η² (grade)'),
]
make_table(['ID', 'Vị trí', 'Vấn đề', 'Ownership', 'Action'], t2_data)


H2('TIER 3 — MEDIUM PRIORITY (3 issues): Optional nhưng nên có')
P('Hoàn thiện chất lượng paper. Reviewers có thể yêu cầu.', italic=True)

t3_data = [
    ('Q1-10', 'Q1 Discussion 4.4', 'Cultural collectivism + developmental task thiếu citations',
     'Em', 'Add: Triandis 1995; Markus & Kitayama 1991; Allen et al. 2013'),
    ('Q3-7', 'Q3 Results', 'Thiếu summary visualization',
     'Em', 'Add Table 6 + Figure 1 (grade trajectory plot)'),
    ('Q3-8', 'Q3 Discussion 4.3', 'Gender brief — chưa justify',
     'Em', 'Add: "(see companion analysis [Cong et al., in prep])"'),
]
make_table(['ID', 'Vị trí', 'Vấn đề', 'Ownership', 'Action'], t3_data)


d.add_page_break()


# ============================================================
H1('ACTION PLAN — TIMELINE 2 TUẦN')

H2('Tuần 1 (01-07/06): Parallel execution')

H2('A. NCS + Thầy actions (4 BLOCKING decisions)')

t_action = [
    ('1', 'NCS confirm qualitative interviews data',
     'NCS', '3 ngày',
     'Số n participants + recruitment + transcripts status'),
    ('2', 'Thầy + NCS decide R² strategy (Q1-8)',
     'Thầy + NCS', '3 ngày',
     '(A) re-run integrated SEM HAY (B) modify claim'),
    ('3', 'NCS xin IRB letter HNUE',
     'NCS', '2-4 tuần (bắt đầu ngay)',
     'Quyết định + ngày + tên hội đồng'),
    ('4', 'Thầy + NCS decide cross-ref strategy Q1↔Q3',
     'Thầy + NCS', '3 ngày',
     '(A) Q1 first → Q3 cite HAY (B) submit cùng lúc'),
]
make_table(['#', 'Action', 'Owner', 'Deadline', 'Output'], t_action)

H2('B. Em parallel fixes (12 HIGH + 3 MEDIUM = 15 issues)')

t_em = [
    ('B1', 'Em build outline v3 Q1 với fixes 1-5, 7, 10', 'Em', '3 ngày',
     'Outline Q1 v3 ~3000 từ'),
    ('B2', 'Em re-extract subtype-specific β table từ LA (Q1-9)', 'Em', '1 ngày',
     'Table 7 predictors × 3 outcomes'),
    ('B3', 'Em build outline v3 Q3 với fixes 1-5, 7-8', 'Em', '2 ngày',
     'Outline Q3 v3 ~1500 từ'),
    ('B4', 'Em verify lại tất cả citations trong v3 vs PDF', 'Em', '1 ngày',
     'Verification log'),
]
make_table(['#', 'Action', 'Owner', 'Deadline', 'Output'], t_em)

H2('Tuần 2 (08-14/06): Integration + Final outline')

t_w2 = [
    ('C1', 'Em integrate decisions từ NCS/thầy vào outline v3',
     'Em', '2 ngày', 'Outline v3 FINAL'),
    ('C2', 'Em add full Ethics statement vào Q3 (Q3-6)', 'Em',
     '0.5 ngày (sau khi có IRB)', 'Q3 Methods 2.5 complete'),
    ('C3', 'Em finalize cross-ref strategy Q1↔Q3 (Q3-9)', 'Em',
     '0.5 ngày', 'Footnote/citation strategy in both papers'),
    ('C4', 'NCS + thầy review outline v3 FINAL', 'NCS + thầy', '2-3 ngày',
     'Approval before draft writing'),
    ('C5', 'Em bắt đầu viết DRAFT Q1', 'Em', 'Tuần 3+', 'Methods → Results → '
     'Discussion → Intro → Abstract'),
]
make_table(['#', 'Action', 'Owner', 'Deadline', 'Output'], t_w2)


d.add_page_break()


# ============================================================
H1('TÓM TẮT — 19 ISSUES PHÂN LOẠI')

t_sum = [
    ('TIER 1 BLOCKING', '4', 'Q1-6, Q1-8, Q3-6, Q3-9', 'NCS + Thầy', 'BLOCKING'),
    ('TIER 2 HIGH', '12', 'Q1-1,2,3,4,5,7,9 + Q3-1,2,3,4,5', 'Em (outline v3)',
     'HIGH'),
    ('TIER 3 MEDIUM', '3', 'Q1-10, Q3-7, Q3-8', 'Em', 'MEDIUM'),
    ('TOTAL', '19', '—', '—', '—'),
]
make_table(['Tier', 'Count', 'Issue IDs', 'Owner', 'Priority'], t_sum)


H2('Em cần thầy + NCS trả lời 4 câu hỏi BLOCKING')

P('Để em build outline v3 hoàn chỉnh, cần answer 4 câu này:', italic=True)

P('1. **Q1-6 Qualitative interviews**: NCS đã phỏng vấn bao nhiêu HS? '
  'Transcripts có sẵn chưa? Sampling strategy thế nào?', italic=False)

P('2. **Q1-8 R² integrated**: Thầy + NCS chọn:', italic=False)
P('   (A) Re-run integrated SEM với 7 predictors đồng thời (cần access bộ data '
  'gốc + 1-2 ngày phân tích), HAY', italic=False)
P('   (B) Modify outline claim — clarify "we report separate-model R² from '
  'parent dissertation" + cite LA', italic=False)

P('3. **Q3-6 IRB letter HNUE**: NCS có letter chính thức chưa? Nếu chưa, '
  'sẵn sàng xin retroactive?', italic=False)

P('4. **Q3-9 Cross-ref strategy**: Thầy + NCS chọn:', italic=False)
P('   (A) Submit Q1 first (BMC Psychiatry) → wait accept → submit Q3 với '
  'citation to Q1 (8-12 tháng total timeline), HAY', italic=False)
P('   (B) Submit Q1 + Q3 cùng lúc với "companion paper under review" note '
  'trong cả 2 (6 tháng timeline, but risk overlap reviewers)', italic=False)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
