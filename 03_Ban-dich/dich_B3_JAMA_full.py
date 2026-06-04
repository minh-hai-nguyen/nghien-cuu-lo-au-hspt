# -*- coding: utf-8 -*-
"""Dịch đầy đủ cụm từ-cụm từ B3 JAMA App CBT 2024"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link bài báo gốc: https://doi.org/10.1001/jamanetworkopen.2024.28372', size=10)
add_p(doc, 'ClinicalTrials.gov: NCT05130281', size=10)

add_heading(doc, 'Hiệu quả can thiệp dựa trên ứng dụng di động cho thanh niên có rối loạn lo âu: Thử nghiệm lâm sàng ngẫu nhiên', 1)
h = doc.add_paragraph()
r = h.add_run('Efficacy of a Mobile App-Based Intervention for Young Adults With Anxiety Disorders: A Randomized Clinical Trial')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Efficacy of a Mobile App-Based Intervention for Young Adults With Anxiety Disorders: A Randomized Clinical Trial'),
    ('Tiêu đề dịch', 'Hiệu quả can thiệp dựa trên ứng dụng di động cho thanh niên có rối loạn lo âu: Thử nghiệm lâm sàng ngẫu nhiên'),
    ('Tác giả', 'Jennifer N. Bress PhD, Avital Falk PhD, Maddy M. Schier BA, Abhishek Jaywant PhD, Elizabeth Moroney PhD, Monika Dargis PhD, Shannon M. Bennett PhD, Matthew A. Scult PhD, Kevin G. Volpp MD PhD, David A. Asch MD MBA, Mohan Balachandran MA MS, Roy H. Perlis MD, Francis S. Lee MD PhD, Faith M. Gunning PhD'),
    ('Cơ quan', 'Weill Cornell Medicine, New York; University of Pennsylvania; Massachusetts General Hospital'),
    ('Tạp chí', 'JAMA Network Open (Q1, IF ≈ 13,8)'),
    ('Xuất bản', '20/08/2024, Vol. 7(8), e2428372, 15 trang'),
    ('DOI', '10.1001/jamanetworkopen.2024.28372'),
    ('Đăng ký thử nghiệm', 'ClinicalTrials.gov NCT05130281'),
    ('Loại NC', 'Thử nghiệm lâm sàng ngẫu nhiên (RCT) — 3 nhóm khuyến khích'),
    ('Mẫu', '59 thanh niên 18–25 tuổi có rối loạn lo âu (GAD 56%, SAD 41%)'),
    ('Can thiệp', 'App Maya — CBT tự hướng dẫn 12 phiên/6 tuần + khuyến khích gamification'),
    ('Đo lường', 'HAM-A (chính), ASI, LSAS, HDRS-24, uMARS'),
    ('Truy cập', 'Open Access — CC-BY'),
])
add_page_ref(doc, '1–15', 'JAMA Network Open', 'Vol. 7(8), 2024')

# TÓM TẮT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Tầm quan trọng: Rối loạn lo âu phổ biến và ít được điều trị ở thanh niên. Các can thiệp sức khỏe tâm thần kỹ thuật số cho lo âu có triển vọng nhưng bị hạn chế bởi phạm vi hẹp các thành phần trị liệu và mức độ tương tác người dùng thấp.')

add_p(doc, 'Mục tiêu: Điều tra hiệu quả và mức tương tác với Maya, một can thiệp liệu pháp nhận thức–hành vi (CBT) di động toàn diện, tự hướng dẫn, có tính mở rộng, với các tính năng tương tác nhúng, so sánh tác động của 3 điều kiện khuyến khích.')

p = doc.add_paragraph()
r = p.add_run('Thiết kế, bối cảnh và người tham gia: Thử nghiệm lâm sàng ngẫu nhiên tuyển thanh niên 18–25 tuổi có rối loạn lo âu qua quảng cáo trực tuyến và phòng khám tâm thần ngoại trú Weill Cornell Medicine. Ghi danh từ 16/06/2021 đến 11/11/2022. Phân tích dữ liệu 21/12/2022–14/06/2024.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('Can thiệp: Người tham gia nhận chương trình 6 tuần của can thiệp và được phân ngẫu nhiên vào 1 trong 3 điều kiện khuyến khích dựa trên tin nhắn: (1) khung mất mát (loss-framed — mất điểm nếu không hoàn thành), (2) khung thu lợi (gain-framed — nhận điểm khi hoàn thành), (3) khung thu lợi + hỗ trợ xã hội (gain–social support — có người thân nhận cập nhật tiến trình).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('Kết quả chính: Kết quả chính là thay đổi triệu chứng lo âu từ ban đầu đến kết thúc điều trị, đo bằng Thang Lo âu Hamilton (HAM-A). Lo âu GIẢM đáng kể qua tất cả các điều kiện: chênh lệch trung bình −5,64 (KTC 95%: −7,23 đến −4,05; Cohen d = 0,94 — LỚN) tại tuần 6, và duy trì tại tuần 12 (chênh lệch −5,67; KTC: −7,29 đến −4,04; Cohen d = 1,04 — RẤT LỚN). KHÔNG có bằng chứng hiệu quả KHÁC NHAU giữa 3 điều kiện khuyến khích. Người tham gia hoàn thành trung bình 10,8/12 phiên (90%). Điểm chất lượng app (uMARS) vượt ngưỡng chấp nhận.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

add_p(doc, 'Kết luận: Giả thuyết chính (khung thu lợi + hỗ trợ xã hội hiệu quả nhất) KHÔNG được ủng hộ. Tuy nhiên, kết quả gợi ý rằng ứng dụng CBT tích hợp đầy đủ kỹ năng CBT và tính năng tương tác nhúng HIỆU QUẢ cải thiện triệu chứng ở thanh niên có rối loạn lo âu, BẤT KỂ loại khuyến khích. Can thiệp kỹ thuật số đại diện bước đi hứa hẹn hướng tới phổ biến rộng rãi can thiệp dựa trên bằng chứng chất lượng cao.')

# TÓM TẮT ĐÁNH GIÁ NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'JAMA Network Open Q1 IF ≈ 13,8 — TẠP CHÍ UY TÍN CỰC CAO. RCT đăng ký (NCT05130281). CONSORT.',
    'App Maya — CBT TỰ HƯỚNG DẪN toàn diện: video tâm lý giáo dục, tái cấu trúc nhận thức, BÀI TẬP PHƠI NHIỄM (exposure — thường bị bỏ qua ở app khác), chánh niệm. 12 phiên/6 tuần.',
    'Cohen d = 0,94 (tuần 6) và 1,04 (tuần 12) — kích thước hiệu ứng LỚN. DUY TRÌ sau 12 tuần.',
    'Hoàn thành 10,8/12 phiên (90%) — mức tương tác CỰC CAO cho app tự hướng dẫn.',
    '98% giữ lại tuần 6, 93% tuần 12 — mất mẫu rất thấp.',
    'Loại khuyến khích KHÔNG ảnh hưởng — app hiệu quả BẤT KỂ gamification.',
    'Thanh niên 18–25, GAD 56% + SAD 41% — phổ rộng rối loạn lo âu.',
    'Phù hợp B2 JMIR (DMHI SAD g=0,508), B9 NMA (iCBT SUCRA 71,2%), B11 Japan iCBT.',
    'Open Access CC-BY.',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'n = 59 — NHỎ (dự kiến 120 nhưng COVID cản trở tuyển mộ). Phân bổ không đều (30/11/18).',
    'Chỉ thanh niên 18–25 — KHÔNG bao gồm VTN <18 (đề tài tập trung cấp 2+3).',
    'KHÔNG có nhóm đối chứng không can thiệp (WL) — chỉ so sánh 3 loại khuyến khích. KHÔNG biết app vs không app.',
    'Chỉ Mỹ, chủ yếu nữ (78%), đa dạng sắc tộc nhưng không đại diện LMIC.',
    'COVID ảnh hưởng tuyển mộ — giảm power.',
    'App Maya cụ thể — có thể không khái quát cho app khác.',
]:
    add_p(doc, f'• {b}')

# 1. GIỚI THIỆU
add_page_ref(doc, '1–3', 'JAMA Network Open', 'Vol. 7(8), 2024')
add_heading(doc, '1. GIỚI THIỆU', 2)

add_p(doc, 'Rối loạn lo âu đã tăng ở người 18–25 tuổi nhanh hơn bất kỳ phạm vi tuổi trưởng thành nào khác, với ước tính tỷ lệ tháng gần đây đạt 15%. Cố vấn Tổng Y sĩ Hoa Kỳ nhấn mạnh rằng mặc dù tốc độ thay đổi tăng nhanh trong đại dịch COVID-19, các chẩn đoán tâm thần ở người trẻ đã tăng từ trước đại dịch. Rối loạn lo âu liên quan đến chất lượng sống kém, khó khăn làm việc, gián đoạn quan hệ xã hội và thành tích học tập, và gánh nặng tài chính.')

add_p(doc, 'Các rối loạn này có thể được điều trị hiệu quả bằng can thiệp tâm lý trị liệu, và đa số bệnh nhân, đặc biệt thanh niên, bày tỏ ưa thích tâm lý trị liệu hơn thuốc hướng thần. Liệu pháp nhận thức–hành vi (CBT) được coi là tiêu chuẩn vàng cho điều trị lo âu dựa trên bằng chứng hiệu quả phong phú. Tuy nhiên, mặc dù tồn tại can thiệp hiệu quả, hơn 25% thanh niên nhận thấy nhu cầu chưa được đáp ứng cho điều trị SKTT, một phần do thách thức tiếp cận: miễn cưỡng thảo luận vấn đề SKTT, ràng buộc tài chính, khoảng cách địa lý từ dịch vụ, và hạn chế đào tạo bác sĩ trong phương pháp CBT.')

add_p(doc, 'Can thiệp kỹ thuật số, như ứng dụng smartphone, cho phép điều trị tại nhà, giá cả phải chăng hơn chăm sóc SKTT truyền thống, và không phụ thuộc vào khả năng tiếp cận nhà trị liệu được đào tạo tại địa phương. Sự tăng tỷ lệ lo âu không được đáp ứng bởi tăng cung nhân lực SKTT, và tiếp cận kỹ thuật số đến CBT tự hướng dẫn dựa trên bằng chứng là giải pháp có tính mở rộng.')

add_p(doc, 'App CBT tự hướng dẫn đã chứng minh hiệu quả ở người lớn trầm cảm, nhưng ít nghiên cứu điều tra hiệu quả ở người có rối loạn lo âu. App SKTT tự hướng dẫn bị cản trở bởi 2 hạn chế: (1) Đa số cung cấp đại diện KHÔNG ĐẦY ĐỦ tâm lý trị liệu — đặc biệt không nhấn mạnh chiến lược hành vi (phơi nhiễm với tình huống gây lo âu), thành phần trị liệu THEN CHỐT cho can thiệp lo âu. (2) Mức tương tác người dùng THẤP giảm tác động tiềm năng.')

add_p(doc, 'App Maya là can thiệp CBT kỹ thuật số nhắm lo âu ở thanh niên giải quyết hạn chế của app hiện có. App dạy đầy đủ kỹ năng CBT nhắm các đặc điểm chung của lo âu, tập trung kỹ năng hành vi thường bị bỏ qua. App sử dụng giao diện hấp dẫn trực quan và tích hợp chiến lược tương tác từ kinh tế hành vi. Nội dung trình bày dạng tương tác khuyến khích tham gia chủ động.')

# 2. PHƯƠNG PHÁP
add_page_ref(doc, '3–6', 'JAMA Network Open', 'Vol. 7(8), 2024')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)

add_p(doc, '2.1. Thiết kế', bold=True)
add_p(doc, 'RCT 3 nhóm khuyến khích. Dự kiến n=120 nhưng chỉ đạt n=59 do COVID. Thanh niên 18–25 có rối loạn lo âu (ADIS-5 CSR ≥4). App 6 tuần, 12 phiên (2 phiên/tuần). Đo tại tuần 0 (ban đầu), tuần 3 (giữa), tuần 6 (kết thúc), tuần 12 (theo dõi). CONSORT.')

add_p(doc, '2.2. Can thiệp — App Maya', bold=True)
add_p(doc, 'Phát triển bởi bác sĩ lâm sàng và nhà nghiên cứu Weill Cornell Medicine. Miễn phí tải. 12 phiên bao gồm: video tâm lý giáo dục + câu đố, bài tập tái cấu trúc nhận thức, BÀI TẬP PHƠI NHIỄM (exposure — thành phần then chốt thường thiếu ở app khác), chánh niệm (mindfulness), và kỹ năng CBT bổ sung. Mỗi phiên bao gồm bài tập về nhà hoàn thành trước phiên tiếp theo.')

add_p(doc, '2.3. Điều kiện khuyến khích', bold=True)
add_p(doc, '(1) Khung mất mát (loss-framed): bắt đầu tuần với số điểm nhất định, MẤT nếu không hoàn thành phiên đúng hạn.')
add_p(doc, '(2) Khung thu lợi (gain-framed): NHẬN điểm khi hoàn thành phiên đúng hạn.')
add_p(doc, '(3) Thu lợi + hỗ trợ xã hội (gain–social support): như (2) + chỉ định bạn/người thân nhận cập nhật tiến trình và gửi tin nhắn hỗ trợ.')
add_p(doc, 'Khuyến khích qua tin nhắn text (Way to Health platform, ĐH Pennsylvania). Huy chương ảo (bronze, silver, gold, platinum). Nội dung và format app KHÔNG khác nhau giữa 3 điều kiện.')

add_p(doc, '2.4. Phân ngẫu nhiên', bold=True)
add_p(doc, 'Way to Health tạo chuỗi phân bổ ngẫu nhiên, 1:1:1 không theo khối. Trợ lý NC bị che giấu phân bổ. Do COVID + không theo khối → phân bổ không đều (30/11/18).')

add_p(doc, '2.5. Thang đo', bold=True)
add_p(doc, '• HAM-A (Hamilton Anxiety Rating Scale, 14 mục) — kết quả CHÍNH. Đo triệu chứng lo âu. Điểm cao = lo âu nặng.')
add_p(doc, '• ASI (Anxiety Sensitivity Index) — lo âu nhạy cảm. Kết quả PHỤ.')
add_p(doc, '• LSAS (Liebowitz Social Anxiety Scale) — lo âu xã hội. Kết quả PHỤ.')
add_p(doc, '• HDRS-24 (Hamilton Depression Rating Scale) — trầm cảm. Kết quả KHÁM PHÁ.')
add_p(doc, '• uMARS (User Mobile Application Rating Scale) — hài lòng app. Ngưỡng ≥4 = chấp nhận.')
add_p(doc, '• ADIS-5 (Anxiety and Related Disorders Interview Schedule for DSM-5) — chẩn đoán. CSR ≥4 = có rối loạn.')

add_p(doc, '2.6. Phân tích', bold=True)
add_p(doc, 'Mô hình hỗn hợp tuyến tính (LMM), random effect = subject intercept, ITT. Tác động cố định: thời gian × điều kiện khuyến khích. Kruskal-Wallis cho engagement (mẫu không đều). Cohen d: 0,2 nhỏ, 0,5 trung bình, 0,8 lớn. Jamovi 2.3.2.0 (R). P < 0,05.')

# 3. KẾT QUẢ
add_page_ref(doc, '6–10', 'JAMA Network Open', 'Vol. 7(8), 2024')
add_heading(doc, '3. KẾT QUẢ', 2)

add_p(doc, '3.1. Đặc điểm mẫu', bold=True)
add_p(doc, '68 ứng viên → 9 loại trừ → 59 ngẫu nhiên hóa: 30 gain + 11 loss + 18 social support. Tuổi TB 23,1 (SD=1,9). Nữ 78% (46). Châu Á 37%, Trắng 42%, Da đen 5%, Hispanic 8%. Đại học 54%, sau ĐH 20%. GAD phổ biến nhất (56%), SAD (41%). HAM-A ban đầu TB 15,0 (SD=6,5). Không có sự cố bất lợi nghiêm trọng.')

add_p(doc, '3.2. Hiệu quả trên lo âu', bold=True)

add_heading(doc, 'Bảng 1. Thay đổi HAM-A (lo âu) theo thời gian — Toàn mẫu', 3)
add_table(doc,
    ['Thời điểm', 'Chênh lệch TB', 'KTC 95%', 'Cohen d', 'p'],
    [['Tuần 3 (giữa)', '−3,20', '−4,76 đến −1,64', '0,64 (TB)', '<0,001'],
     ['Tuần 6 (kết thúc)', '−5,64', '−7,23 đến −4,05', '0,94 (LỚN)', '<0,001'],
     ['Tuần 12 (theo dõi)', '−5,67', '−7,29 đến −4,04', '1,04 (RẤT LỚN)', '<0,001']],
    widths=[3.0, 2.5, 3.5, 2.5, 1.5])
add_p(doc, 'Lo âu GIẢM LIÊN TỤC từ tuần 0→3→6, DUY TRÌ tại tuần 12. Kích thước hiệu ứng LỚN (d=0,94) và RẤT LỚN (d=1,04) — vượt trội so với B2 JMIR DMHI (g=0,508).', size=9, italic=True)

add_heading(doc, 'Bảng 2. Kết quả PHỤ — ASI và LSAS', 3)
add_table(doc,
    ['Thang đo', 'Tuần 6 chênh lệch', 'd', 'Tuần 12 chênh lệch', 'd'],
    [['ASI (lo âu nhạy cảm)', '−9,51', '0,93', '−8,90', '0,93'],
     ['LSAS (lo âu XH)', '−12,67', '0,70', '−17,61', '1,07']],
    widths=[3.5, 3.0, 1.5, 3.0, 1.5])
add_p(doc, 'LSAS (lo âu XH) cải thiện LỚN NHẤT tại follow-up (d=1,07) — app hiệu quả cho cả GAD lẫn SAD.', size=9, italic=True)

add_p(doc, '3.3. So sánh khuyến khích', bold=True)
add_p(doc, 'KHÔNG có bằng chứng chênh lệch giữa 3 điều kiện khuyến khích cho bất kỳ thang đo nào (tất cả p > 0,05). Loss-framed vs Social support: chênh lệch −1,40 (KTC: −4,72 đến 1,93). Gain-framed vs Social support: 1,38 (KTC: −1,19 đến 3,96). Nghĩa là: app hiệu quả BẤT KỂ loại gamification.')

add_p(doc, '3.4. Tương tác (Engagement)', bold=True)
add_p(doc, 'Trung bình hoàn thành 10,8/12 phiên (90%). 64% hoàn thành TẤT CẢ 12 phiên. 98% giữ lại tuần 6. 93% giữ lại tuần 12. 39% tiếp tục dùng app sau 6 tuần. Số phiên hoàn thành KHÔNG khác biệt theo tuổi, giới, sắc tộc, hoặc điều kiện khuyến khích.')

add_p(doc, '3.5. Hài lòng', bold=True)
add_p(doc, 'Điểm uMARS App Quality VƯỢT ngưỡng chấp nhận (≥4) tại tất cả thời điểm. Không khác biệt theo thời gian hoặc điều kiện khuyến khích.')

# 4. THẢO LUẬN
add_page_ref(doc, '10–13', 'JAMA Network Open', 'Vol. 7(8), 2024')
add_heading(doc, '4. THẢO LUẬN', 2)

add_p(doc, 'Nghiên cứu này đánh giá hiệu quả app CBT di động tự hướng dẫn (Maya) cho thanh niên có rối loạn lo âu. Giả thuyết chính — khuyến khích thu lợi + hỗ trợ xã hội hiệu quả nhất — KHÔNG được ủng hộ. Tuy nhiên, kết quả cho thấy giảm LỚN và đáng kể trong lo âu qua tất cả các điều kiện, với kích thước hiệu ứng lớn (d=0,94) và duy trì sau 12 tuần.')

add_p(doc, 'So sánh với NC trước: kích thước hiệu ứng (d=0,94) VƯỢT TRỘI so với đa số NC app CBT cho trầm cảm (d trung bình ~0,3–0,5). Lý do có thể: (1) Maya bao gồm BÀI TẬP PHƠI NHIỄM — thành phần then chốt thường bị bỏ qua ở app khác; (2) giao diện tương tác khuyến khích tham gia chủ động; (3) gamification nhúng tăng tương tác.')

add_p(doc, 'Mức tương tác (90% hoàn thành) vượt trội so với app SKTT khác (thường 50–70%). 64% hoàn thành TẤT CẢ phiên — rất cao. 39% tiếp tục dùng sau chương trình — gợi ý giá trị cảm nhận.')

add_p(doc, 'Không có chênh lệch khuyến khích gợi ý: app với tính năng tương tác TỐT (nội dung CBT đầy đủ + giao diện hấp dẫn) có thể KHÔNG CẦN thêm gamification phức tạp. Điều này đơn giản hóa triển khai.')

add_p(doc, 'Hạn chế:', bold=True)
add_p(doc, '(1) n=59 nhỏ hơn mục tiêu (120) do COVID. Giảm power và tạo phân bổ không đều.')
add_p(doc, '(2) Không có nhóm đối chứng không can thiệp (WL) — không biết cải thiện do app hay do thời gian/regression to mean.')
add_p(doc, '(3) Chỉ thanh niên 18–25 — cần NC ở VTN <18.')
add_p(doc, '(4) Đánh giá viên không hoàn toàn mù — có thể thiên lệch.')
add_p(doc, '(5) Chỉ Mỹ, chủ yếu nữ (78%), đại học+ — không đại diện đa dạng.')

# 5. KẾT LUẬN
add_heading(doc, '5. KẾT LUẬN', 2)
add_p(doc, 'Trong thử nghiệm lâm sàng ngẫu nhiên này về can thiệp dựa trên app cho lo âu, giả thuyết chính rằng cải thiện lo âu sẽ lớn nhất trong điều kiện thu lợi + hỗ trợ xã hội KHÔNG được ủng hộ. Tuy nhiên, kết quả gợi ý rằng app CBT di động tự hướng dẫn tích hợp đầy đủ kỹ năng CBT và tính năng tương tác HIỆU QUẢ cải thiện triệu chứng ở thanh niên có rối loạn lo âu (Cohen d = 0,94), với mức tương tác cao (90% hoàn thành) và duy trì sau 12 tuần (d = 1,04). Can thiệp kỹ thuật số đại diện bước đi hứa hẹn hướng tới phổ biến rộng rãi can thiệp chất lượng cao.')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
for ref in [
    'Bress, J.N. et al. (2024). Efficacy of a Mobile App-Based Intervention for Young Adults With Anxiety Disorders. JAMA Network Open, 7(8), e2428372.',
    'Hamilton, M. (1959). The assessment of anxiety states by rating. British Journal of Medical Psychology, 32(1), 50–55.',
    'Hofmann, S.G. et al. (2012). The efficacy of CBT: a review of meta-analyses. Cognitive Therapy and Research, 36, 427–440.',
    '(Xem đầy đủ 56 tài liệu tham khảo trong bài gốc)',
]:
    add_p(doc, ref, size=10)

# VIẾT TẮT
add_abbreviation_table(doc, [
    ('CBT', 'Cognitive Behavioral Therapy — Liệu pháp Nhận thức–Hành vi'),
    ('HAM-A', 'Hamilton Anxiety Rating Scale — Thang Lo âu Hamilton'),
    ('ASI', 'Anxiety Sensitivity Index — Chỉ số Nhạy cảm Lo âu'),
    ('LSAS', 'Liebowitz Social Anxiety Scale — Thang Lo âu Xã hội Liebowitz'),
    ('HDRS', 'Hamilton Depression Rating Scale'),
    ('uMARS', 'User Mobile Application Rating Scale'),
    ('ADIS-5', 'Anxiety and Related Disorders Interview Schedule for DSM-5'),
    ('GAD', 'Generalized Anxiety Disorder — Rối loạn Lo âu Tổng quát'),
    ('SAD', 'Social Anxiety Disorder — Rối loạn Lo âu Xã hội'),
    ('LMM', 'Linear Mixed-effects Model — Mô hình Hỗn hợp Tuyến tính'),
    ('ITT', 'Intention-to-Treat — Phân tích Ý định Điều trị'),
    ('RCT', 'Randomized Clinical Trial'),
    ('CONSORT', 'Consolidated Standards of Reporting Trials'),
])

# PHẢN BIỆN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')

add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'JAMA Network Open Q1 IF ≈ 13,8 — tạp chí UY TÍN CỰC CAO. RCT đăng ký NCT05130281. CONSORT.',
    'App Maya bao gồm BÀI TẬP PHƠI NHIỄM — thành phần THEN CHỐT thường bị bỏ qua ở app khác (B2 JMIR nhấn mạnh). AJP QT28 (Zugman): exposure là core CBT.',
    'Cohen d = 0,94 (tuần 6) và 1,04 (tuần 12) — kích thước hiệu ứng LỚN + DUY TRÌ. Vượt trội B2 JMIR (DMHI g=0,508 tổng) và B9 NMA (iCBT SUCRA 71,2%).',
    'Tương tác 90% (10,8/12 phiên) — CỰC CAO cho app tự hướng dẫn. 64% hoàn thành TẤT CẢ.',
    'Không cần khuyến khích phức tạp — app TỰ NÓ đủ tốt. Đơn giản hóa triển khai.',
    'Đa dạng sắc tộc (37% Á, 42% Trắng, 5% Đen, 8% Hispanic) — đại diện tốt hơn đa số NC.',
    'ASI + LSAS giảm d > 0,7–1,07 — app hiệu quả cho CẢ GAD + SAD.',
    'Open Access CC-BY.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'n = 59 NHỎ (dự kiến 120) — COVID cản trở. Phân bổ không đều 30/11/18. Giảm power đáng kể.',
    'KHÔNG CÓ nhóm WL/đối chứng không can thiệp — KHÔNG biết cải thiện do app hay tự nhiên. So: B8 Sri Lanka RCT CÓ đối chứng → kết luận mạnh hơn.',
    'Chỉ 18–25 tuổi — KHÔNG bao gồm VTN <18 (cấp 2+3 của đề tài). B8 Sri Lanka: HS lớp 9 (~14 tuổi). B9 NMA: trẻ + VTN.',
    'Chỉ Mỹ (Weill Cornell Medicine, NYC) — không đại diện LMIC. B7 CA-CBT ĐNA: 0 NC từ VN.',
    'Nữ 78% — thiên lệch giới. Đa số ĐH+ — thiên lệch học vấn.',
    'App Maya cụ thể — kết quả có thể không khái quát cho app CBT khác. Tuy nhiên, app MIỄN PHÍ + bao gồm exposure → mô hình tốt.',
    'Không đo chi phí-hiệu quả — quan trọng cho LMIC/VN.',
    'Tuyển qua quảng cáo + phòng khám — thiên lệch chọn (chỉ người tìm kiếm giúp đỡ).',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'VN CẦN phát triển app CBT tiếng Việt tương tự Maya — bao gồm PHƠI NHIỄM + tương tác. Miễn phí hoặc chi phí thấp.',
    'RCT app CBT cho VTN <18 tại VN — mẫu cho đề cương giai đoạn 2. Kết hợp B3 (app) + B8 (GV hướng dẫn) + B7 (thích ứng VH).',
    'So sánh: app tự hướng dẫn vs app có GV/mentor hướng dẫn tại VN — B2 JMIR: có hướng dẫn g=0,825 > không hướng dẫn ~0,3–0,4.',
    'Đo chi phí-hiệu quả app vs CBT mặt đối mặt tại VN — xác định giải pháp kinh tế nhất.',
    'Kiểm tra app ở VTN có SAD riêng (Jefferies QT35: VN SAD = 30,7%) — B9 NMA: iCBT hạng 1 cho SAD.',
    'Tích hợp exposure (phơi nhiễm) vào app VN — thành phần THEN CHỐT mà app khác thường bỏ qua.',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '49_JAMA_App_CBT_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['59', '23,1', '0,94', '1,04', '10,8', '15,0', 'HAM-A', 'Maya', 'CONSORT', '−5,64', '−5,67', '78%', '98%']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
