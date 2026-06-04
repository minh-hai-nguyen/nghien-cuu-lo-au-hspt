"""Build doc: 'VN014 H.T. Học - Khái niệm "quan hệ cha mẹ-con" rộng quá' (CTH v6 + CÂU TRẢ LỜI tô xanh).
Verified facts:
- VN014: β = 0.272 cho 'tính chất mối quan hệ với cha mẹ' (Tom-tat VN014 confirmed)
- VN003 Pham 2024: chăm sóc cảm xúc β = -0.40, P<0.001 (Tom-tat VN003 confirmed)
- QT009 Qiu 2022: EMBU 9 chiều (bản TQ), nuôi dạy tích cực OR=0.30/0.32
- EMBU Perris 1980 bản gốc: 81 mục, 15 tiểu thang (verified WebSearch)
- s-EMBU rút gọn: 3 nhóm (rejection, emotional warmth, overprotection)
- PARQ Rohner: 60 mục, 4 chiều (warmth, hostility, indifference, undifferentiated rejection) verified
- Compas 2017: lý luận về biến đa chiều — không gộp
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Phan_bien_VN014_quan_he_cha_me_rong_qua.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE
    return p

def blue_bullet(text, size=12):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = BLUE

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẢN BIỆN VN014 HOÀNG TRUNG HỌC (2025)\nKHÁI NIỆM "QUAN HỆ CHA MẸ-CON" CÓ RỘNG QUÁ KHÔNG?')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY

para('')

# Câu hỏi tô xanh
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run('Trong bài của ông Học thì quan hệ cha mẹ có rộng quá không?', italic=True)

# 1. Số liệu cụ thể VN014
H('1. Khái niệm "quan hệ cha mẹ-con" trong VN014', level=2, color=NAVY)
para(
    'Hoàng Trung Học và cộng sự (2025) trên 8.473 học sinh THCS-THPT '
    'tại 6 tỉnh Việt Nam (Đợt 1 trong COVID n = 4.052 / Đợt 2 sau '
    'COVID n = 4.337) sử dụng DASS-21 đo lo âu-trầm cảm-căng thẳng '
    'kết hợp BẢNG HỎI BỔ SUNG về giấc ngủ, sử dụng thiết bị điện tử, '
    'hoạt động thể chất, và "quan hệ gia đình".', indent=False
)
para(
    'Cụ thể về biến quan hệ cha mẹ-con trong báo cáo:'
)
add_table(
    ['Đặc điểm', 'VN014 H.T. Học (2025)'],
    [
        ['Tên biến', '"Tính chất mối quan hệ với cha mẹ"'],
        ['Hệ số ảnh hưởng', 'β = 0,272 (mạnh nhất trong mô hình)'],
        ['Thang đo', 'Bảng hỏi bổ sung tự thiết kế (không nêu rõ thang chuẩn)'],
        ['Số mục', 'Không công bố cụ thể'],
        ['Số chiều phân biệt', 'CHỈ MỘT — gộp tất cả vào 1 biến'],
        ['Phong cách nuôi dạy', 'Không phân biệt'],
        ['Kiểm soát tâm lý vs hành vi', 'Không phân biệt'],
        ['Sự ấm áp vs mâu thuẫn', 'Không phân biệt'],
        ['Giao tiếp / gắn bó', 'Không phân biệt'],
    ]
)

para('')
para(
    'Như vậy, β = 0,272 trong VN014 đại diện cho TÁC ĐỘNG TỔNG HỢP '
    'của TẤT CẢ các thành phần của quan hệ cha mẹ-con — KHÔNG biết '
    'thành phần NÀO mới là yếu tố BẢO VỆ thực sự, thành phần NÀO mới '
    'là yếu tố NGUY CƠ.'
)

# 2. So sánh với chuẩn quốc tế
H('2. So sánh với chuẩn quốc tế — VN014 đo TÁC ĐỘNG TỔNG, không phân tách', level=2, color=NAVY)
para(
    'Y văn quốc tế đã từ lâu phân biệt nhiều CHIỀU CON của quan hệ '
    'cha mẹ-con. Bảng dưới đây tổng hợp năm thang đo chuẩn được sử '
    'dụng nhiều nhất.', indent=False
)
caption('Bảng 1. Năm thang đo chuẩn quốc tế cho quan hệ cha mẹ-con')
add_table(
    ['Thang đo', 'Tác giả gốc', 'Số mục', 'Số chiều'],
    [
        ['EMBU (My Memories of Upbringing)',
         'Perris và cộng sự (1980)', '81', '15 tiểu thang'],
        ['s-EMBU (rút gọn)',
         'Arrindell và cộng sự (1999)', '23–69', '3 (rejection / emotional warmth / overprotection)'],
        ['EMBU phiên bản 9 chiều',
         'Bản Trung Quốc — Qiu 2022 dùng', '~58', '9'],
        ['PARQ (Parental Acceptance-Rejection Q.)',
         'Rohner (1976/1980); bản trưởng thành', '60', '4 (warmth / hostility / indifference / undifferentiated rejection)'],
        ['PSDQ (Parenting Styles & Dimensions)',
         'Robinson và cộng sự (2001)', '32', '3 (authoritative / authoritarian / permissive)'],
        ['PBI (Parental Bonding Instrument)',
         'Parker và cộng sự (1979)', '25', '2 (care / overprotection)'],
        ['CRPBI (Children\'s Report of Parental Behavior)',
         'Schaefer (1965); bản rút gọn', '30', '3 (acceptance / control / autonomy)'],
    ]
)
para('')
para(
    'Trái với VN014, các nghiên cứu quốc tế và một số nghiên cứu Việt '
    'Nam đã PHÂN TÁCH chiều con của quan hệ cha mẹ — và phát hiện hệ '
    'số tác động riêng cho từng chiều. Hai ví dụ tiêu biểu:'
)
para(
    'Thứ nhất, Phạm và cộng sự (2024) — VN003 trong cơ sở dữ liệu — '
    'khảo sát 273 + 273 thanh thiếu niên tại các cơ sở hỗ trợ xã hội '
    'Huế đã PHÂN TÁCH chăm sóc CẢM XÚC khỏi chăm sóc THỂ CHẤT. Kết '
    'quả: chăm sóc cảm xúc β = −0,40 (P < 0,001) — bảo vệ MẠNH cho '
    'sức khỏe tâm thần; chăm sóc thể chất KHÔNG có tác động ý nghĩa. '
    'Nói cách khác, hai chiều con của "chăm sóc" có tác động TRÁI '
    'ngược nhau về cường độ — gộp chung sẽ làm mờ phát hiện.'
)
para(
    'Thứ hai, Qiu và cộng sự (2022) — QT009 trong cơ sở dữ liệu — '
    'trên 2.079 học sinh THCS Trung Quốc sử dụng EMBU bản 9 chiều, '
    'sau đó áp dụng phân tích hồ sơ tiềm ẩn (latent profile analysis) '
    'để xác định BA HỒ SƠ NUÔI DẠY: tích cực (58,6%), trung bình '
    '(32,2%), tiêu cực (9,1%). Hồ sơ tích cực giảm nguy cơ trầm cảm '
    '70% (OR = 0,30) và lo âu 68% (OR = 0,32) so với hồ sơ tiêu '
    'cực. Phù hợp với nguyên tắc PHÂN BIỆT chiều con của ứng phó '
    'do Compas và cộng sự (2017) thiết lập trên 80.850 trẻ em + vị '
    'thành niên.'
)

# 3. Sáu chiều con bị gộp chung
H('3. Sáu chiều con của "quan hệ cha mẹ-con" bị VN014 gộp chung', level=2, color=NAVY)
para(
    'Y văn quốc tế phân biệt ít nhất sáu chiều con — TRÁI ngược nhau '
    'về tác động lên lo âu vị thành niên:', indent=False
)
caption('Bảng 2. Sáu chiều con của "quan hệ cha mẹ-con" và tác động dự kiến với lo âu')
add_table(
    ['#', 'Chiều con', 'Tác động dự kiến với lo âu', 'Bằng chứng tham chiếu'],
    [
        ['1', 'Sự ấm áp / chăm sóc cảm xúc (warmth, emotional care)',
         'β ÂM mạnh — bảo vệ',
         'Phạm 2024 β = −0,40; Qiu 2022 OR = 0,32'],
        ['2', 'Kiểm soát tâm lý (psychological control: gây tội lỗi, rút tình cảm)',
         'β DƯƠNG — yếu tố nguy cơ',
         'Barber 1996; Soenens & Vansteenkiste 2010'],
        ['3', 'Kiểm soát hành vi (behavioral control: giám sát, đặt quy tắc)',
         'β trung tính / ÂM nhẹ',
         'Pinquart 2017'],
        ['4', 'Mâu thuẫn cha mẹ-con (parent-child conflict)',
         'β DƯƠNG mạnh — nguy cơ',
         'Smetana 2011'],
        ['5', 'Giao tiếp cha mẹ-con cởi mở (parent-child communication)',
         'β ÂM — bảo vệ',
         'Liu 2024'],
        ['6', 'Gắn bó / sự gần gũi (attachment, closeness)',
         'β ÂM — bảo vệ',
         'Bowlby 1969 / Allen 2008'],
    ]
)
para('')
para(
    'Hai luận điểm quan trọng từ bảng trên. Thứ nhất, sáu chiều con '
    'có ĐỘ MẠNH và CHIỀU TÁC ĐỘNG khác nhau — gộp chúng vào 1 biến '
    'TRIỆT TIÊU lẫn nhau. Thứ hai, gộp chung có thể tạo ra hệ số '
    'tổng hợp THẤP HƠN giá trị thực của từng chiều con — ví dụ '
    'β = 0,272 của VN014 có thể "giấu" β = −0,40 hoặc thậm chí '
    'mạnh hơn của riêng "sự ấm áp" — và "giấu" β DƯƠNG của riêng '
    '"mâu thuẫn".'
)

# 4. Phân tích điểm yếu
H('4. Bốn hậu quả phương pháp luận', level=2, color=NAVY)
para('Việc gộp tất cả thành 1 biến gây ra bốn hậu quả nghiêm trọng cho giá trị khoa học của VN014.', indent=False)

H('4.1. Hệ số β = 0,272 KHÔNG phân biệt được aspect bảo vệ vs aspect nguy cơ', level=3)
para(
    'Khi β = 0,272 cho biết "quan hệ cha mẹ tốt thì lo âu giảm" — '
    'nhưng không trả lời được: cha mẹ phải làm GÌ để "tốt"? Tăng '
    'sự ấm áp? Giảm kiểm soát tâm lý? Cải thiện giao tiếp? Giảm '
    'mâu thuẫn? Mỗi can thiệp khác nhau cho mỗi chiều con. Hệ số '
    'tổng không hướng dẫn được thực hành.', indent=False
)

H('4.2. Không so sánh trực tiếp được với y văn quốc tế', level=3)
para(
    'β = 0,272 từ thang tự thiết kế của VN014 KHÔNG so sánh được '
    'với β = −0,40 của Phạm 2024 (chăm sóc cảm xúc), OR = 0,30 '
    'của Qiu 2022 (hồ sơ nuôi dạy tích cực), hay các phát hiện '
    'EMBU/PARQ trong y văn quốc tế. Khoảng cách so sánh này làm '
    'giảm giá trị đóng góp của VN014 cho y văn — bài chỉ "đóng '
    'góp số liệu Việt Nam" mà không cho phép so sánh hệ số.', indent=False
)

H('4.3. Không hướng dẫn được thiết kế can thiệp', level=3)
para(
    'Khuyến nghị "đào tạo kỹ năng cha mẹ" trong VN014 quá chung '
    'chung. Đào tạo cha mẹ về CHIỀU CON nào? Mô hình EACP '
    '(Lochman 2025) đào tạo 16 buổi 90 phút cho cha mẹ — TỪNG '
    'buổi nhắm vào kỹ năng cụ thể (giải quyết xung đột, củng '
    'cố hành vi tích cực, lắng nghe phản chiếu). Không có '
    'phân tách chiều con thì không có can thiệp đích.', indent=False
)

H('4.4. Có thể phóng đại hoặc thu nhỏ tác động thực sự', level=3)
para(
    'Hệ số tổng hợp có rủi ro thiên lệch theo HƯỚNG NGẪU NHIÊN '
    'tùy nội dung của bảng hỏi tự thiết kế. Nếu các mục thiên '
    'về "sự ấm áp" — hệ số sẽ phản ánh chủ yếu chiều này. Nếu '
    'thiên về "mâu thuẫn" — chiều ngược lại. Không có cách kiểm '
    'tra do bảng hỏi không công bố. Đây là vấn đề về độ giá trị '
    '(validity) chứ không chỉ độ chi tiết.', indent=False
)

# 5. Đề xuất
H('5. Đề xuất cho thiết kế nghiên cứu mới', level=2, color=NAVY)
para(
    'Đề tài của thầy có thể LẤP KHOẢNG TRỐNG này bằng cách đo "quan '
    'hệ cha mẹ-con" theo nhiều chiều ngay từ thiết kế.', indent=False
)
para(
    'KHUYẾN NGHỊ 1 — Chọn thang chuẩn quốc tế. Ba lựa chọn ưu '
    'tiên cho học sinh THCS Việt Nam:'
)
add_table(
    ['Lựa chọn', 'Thang', 'Ưu điểm', 'Nhược điểm'],
    [
        ['ƯU TIÊN 1',
         'PARQ trẻ em (Rohner 1976/2005), 60 mục, 4 chiều',
         'Đã được sử dụng tại 60+ quốc gia; có bản trẻ em + bản trưởng thành; α = 0,89',
         'Cần dịch + chuẩn hóa cho VN'],
        ['ƯU TIÊN 2',
         's-EMBU (Arrindell 1999), 23 mục, 3 chiều',
         'Rút gọn, dễ áp dụng; bản tiếng TQ + Tây Ban Nha sẵn',
         'Chỉ 3 chiều — kém chi tiết hơn PARQ'],
        ['ƯU TIÊN 3',
         'PBI (Parker 1979), 25 mục, 2 chiều (care + overprotection)',
         'Ngắn, nhanh; nhiều bài VN đã dùng',
         '2 chiều — bỏ sót giao tiếp, mâu thuẫn'],
    ]
)
para('')
para(
    'KHUYẾN NGHỊ 2 — Phân tích RIÊNG hiệu ứng từng chiều với lo '
    'âu. Dự kiến kết quả phù hợp y văn quốc tế: ấm áp β ÂM mạnh; '
    'kiểm soát tâm lý β DƯƠNG; mâu thuẫn β DƯƠNG; giao tiếp β '
    'ÂM. Nếu kết quả khớp dự kiến, đây là bằng chứng VIỆT NAM '
    'bổ sung cho phát hiện quốc tế của Phạm 2024 + Qiu 2022.'
)
para(
    'KHUYẾN NGHỊ 3 — Áp dụng phân tích hồ sơ tiềm ẩn (LPA) như '
    'Qiu và cộng sự (2022). Sau khi đo PARQ/EMBU, áp dụng LPA '
    'để xác định 3–4 hồ sơ nuôi dạy phổ biến trong dân số học '
    'sinh THCS Việt Nam. So sánh tỷ lệ hồ sơ tích cực-trung '
    'bình-tiêu cực với Trung Quốc Qiu 2022 (58,6% / 32,2% / '
    '9,1%) — kỳ vọng tỷ lệ Việt Nam tương đương hoặc thấp hơn '
    'do tác động COVID + áp lực học tập.'
)
para(
    'KHUYẾN NGHỊ 4 — Thiết kế can thiệp ĐÍCH theo chiều con. '
    'Theo mô hình EACP (Lochman 2025): 16 buổi cha mẹ — chia '
    'theo nhóm chiều con cần cải thiện. Cha mẹ có "ấm áp '
    'thấp" → buổi tăng sự ấm áp; cha mẹ "kiểm soát tâm lý '
    'cao" → buổi giảm kiểm soát; cha mẹ "mâu thuẫn cao" → '
    'buổi giải quyết xung đột. Khác biệt với approach chung '
    'chung của VN014.'
)

# 6. CÂU TRẢ LỜI tô xanh
H('6. CÂU TRẢ LỜI', level=2, color=NAVY)
blue_run('Tóm gọn:', bold=True)
blue_bullet(
    'CÓ — khái niệm "quan hệ cha mẹ-con" trong VN014 RỘNG QUÁ. '
    'Tác giả gộp tất cả thành phần (sự ấm áp, kiểm soát tâm lý, '
    'kiểm soát hành vi, mâu thuẫn, giao tiếp, gắn bó) vào MỘT '
    'biến với β = 0,272 — không phân biệt được thành phần nào '
    'bảo vệ, thành phần nào nguy cơ.'
)
blue_bullet(
    'Trái với VN014, Phạm và cộng sự (2024) đã tách chăm sóc '
    'CẢM XÚC khỏi chăm sóc THỂ CHẤT — phát hiện cảm xúc β = '
    '−0,40 (P < 0,001) trong khi thể chất KHÔNG có ý nghĩa. '
    'Qiu và cộng sự (2022) dùng EMBU 9 chiều + LPA — xác định '
    'ba hồ sơ nuôi dạy với OR khác nhau rõ rệt.'
)
blue_bullet(
    'Y văn quốc tế phân biệt ít nhất 6 chiều con: ấm áp / kiểm '
    'soát tâm lý / kiểm soát hành vi / mâu thuẫn / giao tiếp / '
    'gắn bó. Sáu chiều có tác động TRÁI ngược nhau lên lo âu — '
    'gộp chung TRIỆT TIÊU phát hiện.'
)
blue_bullet(
    'Bốn hậu quả của việc gộp: (1) không phân biệt được aspect '
    'bảo vệ vs nguy cơ; (2) không so sánh được với y văn quốc '
    'tế; (3) không hướng dẫn được thiết kế can thiệp đích; '
    '(4) hệ số có thể thiên lệch tùy nội dung bảng hỏi không '
    'công bố.'
)
blue_bullet(
    'Đề xuất cho đề tài của thầy: dùng PARQ (Rohner, 60 mục, '
    '4 chiều — ưu tiên 1) hoặc s-EMBU (23 mục, 3 chiều — ưu '
    'tiên 2) hoặc PBI (25 mục, 2 chiều — ưu tiên 3); phân '
    'tích RIÊNG từng chiều với lo âu; áp dụng LPA để xác '
    'định hồ sơ nuôi dạy.'
)

blue_run('Cách trích vào báo cáo CTH:', bold=True)
blue_run(
    '"Hoàng Trung Học và cộng sự (2025) khảo sát 8.473 học sinh '
    'tại 6 tỉnh Việt Nam, phát hiện quan hệ cha mẹ-con là yếu '
    'tố mạnh nhất ảnh hưởng lo âu (β = 0,272). Tuy nhiên, biến '
    '"quan hệ cha mẹ-con" trong nghiên cứu được đo bằng bảng '
    'hỏi tự thiết kế chỉ một biến tổng hợp — không phân biệt '
    'các chiều con như sự ấm áp, kiểm soát tâm lý, mâu thuẫn, '
    'hay giao tiếp. Trái với VN014, Phạm và cộng sự (2024) tại '
    'Huế đã phân tách chăm sóc cảm xúc khỏi chăm sóc thể chất '
    '— phát hiện cảm xúc β = −0,40 (P < 0,001) trong khi thể '
    'chất KHÔNG có ý nghĩa thống kê. Tương tự, Qiu và cộng sự '
    '(2022) tại Trung Quốc dùng thang EMBU đa chiều kết hợp '
    'phân tích hồ sơ tiềm ẩn để xác định ba hồ sơ nuôi dạy '
    'với OR khác nhau rõ rệt. Khuyến nghị nghiên cứu tiếp '
    'theo tại Việt Nam sử dụng các thang chuẩn quốc tế đa '
    'chiều — như PARQ (Rohner, 60 mục, 4 chiều) hoặc s-EMBU '
    '(Arrindell, 23 mục, 3 chiều) — kết hợp phân tích riêng '
    'từng chiều con với lo âu để xác định thành phần bảo vệ '
    'cụ thể, tạo cơ sở cho can thiệp đích."',
    italic=True
)

# 7. Phụ lục TLTK
H('7. Phụ lục — Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Arrindell, W. A., Sanavio, E., Aguilar, G., Sica, C., Hatzichristou, C., Eisemann, M., Recinos, L. A., Gaszner, P., Peter, M., Battagliese, G., Kállai, J., & van der Ende, J. (1999). The development of a short form of the EMBU: Its appraisal with students in Greece, Guatemala, Hungary and Italy. Personality and Individual Differences, 27(4), 613–628.',
    'Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991. https://doi.org/10.1037/bul0000110',
    'Hoàng, T. H., và cộng sự. (2025). Sức khỏe tâm thần của học sinh Việt Nam trong và sau đại dịch COVID-19. [VN014 trong cơ sở dữ liệu dự án.]',
    'Lochman, J. E., và cộng sự. (2025). Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology.',
    'Parker, G., Tupling, H., & Brown, L. B. (1979). A parental bonding instrument. British Journal of Medical Psychology, 52(1), 1–10.',
    'Perris, C., Jacobsson, L., Lindström, H., von Knorring, L., & Perris, H. (1980). Development of a new inventory for assessing memories of parental rearing behaviour. Acta Psychiatrica Scandinavica, 61(4), 265–274. https://doi.org/10.1111/j.1600-0447.1980.tb00581.x',
    'Phạm, ..., và cộng sự. (2024). Mối quan hệ giữa chăm sóc xã hội và sức khỏe tâm thần ở thanh thiếu niên tại các cơ sở hỗ trợ xã hội tại Huế, Việt Nam. [VN003 trong cơ sở dữ liệu dự án.]',
    'Qiu, Z., Guo, Y., Wang, J., & Zhang, H. (2022). Associations of parenting style and resilience with depression and anxiety symptoms among Chinese middle school students. Frontiers in Public Health. [QT009 trong cơ sở dữ liệu dự án.]',
    'Robinson, C. C., Mandleco, B., Olsen, S. F., & Hart, C. H. (2001). The Parenting Styles and Dimensions Questionnaire (PSDQ). In B. F. Perlmutter, J. Touliatos, & G. W. Holden (Eds.), Handbook of family measurement techniques: Vol. 3 (pp. 319–321). Sage.',
    'Rohner, R. P. (2005). Parental Acceptance-Rejection Questionnaire (PARQ): Test manual. In R. P. Rohner & A. Khaleque (Eds.), Handbook for the study of parental acceptance and rejection (4th ed., pp. 43–106). Rohner Research Publications.',
    'Schaefer, E. S. (1965). Children\'s reports of parental behavior: An inventory. Child Development, 36(2), 413–424.',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
