# -*- coding: utf-8 -*-
"""Tao 2 file QD Hoi dong Dao duc:
- Ban MAU (bo trong [...])
- Ban DIEN SAN (cho NCS Cong Thi Hang research)
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

def build_qd(out_path, is_blank=True):
    """Build mot ban QD phe duyet (blank template hoac dien san)."""
    d = Document()
    for sec in d.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.line_spacing = 1.5

    # Field substitution dict
    if is_blank:
        F = {
            'TEN_BO_NGANH': '[TÊN BỘ/NGÀNH CHỦ QUẢN]',
            'TEN_TRUONG': '[TÊN TRƯỜNG/VIỆN]',
            'TEN_HOI_DONG': '[Tên Hội đồng Đạo đức]',
            'SO_QD': '[___/QĐ-…]',
            'DIA_DIEM': '[Địa điểm]',
            'NGAY': '[ngày]', 'THANG': '[tháng]', 'NAM': '[năm]',
            'TEN_DE_TAI': '[Tên đầy đủ của đề tài/dự án nghiên cứu]',
            'TEN_NCS': '[Họ và tên Chủ nhiệm đề tài/Nghiên cứu sinh]',
            'DON_VI': '[Đơn vị công tác/học tập]',
            'HUONG_DAN': '[Họ tên + học hàm/học vị của Người hướng dẫn]',
            'MUC_TIEU': '[Mục tiêu nghiên cứu — tóm tắt 2-3 câu]',
            'DOI_TUONG': '[Đối tượng nghiên cứu]',
            'CO_MAU': '[Cỡ mẫu dự kiến]',
            'DIA_BAN': '[Địa bàn nghiên cứu]',
            'THOI_GIAN_BD': '[tháng/năm bắt đầu]',
            'THOI_GIAN_KT': '[tháng/năm kết thúc]',
            'PHAM_VI_CONG_BO': '[Mô tả phạm vi công bố: luận án + các bài báo dự '
                               'kiến công bố trên tạp chí trong nước/quốc tế]',
            'CHU_TICH_HD': '[Họ tên + học hàm/học vị Chủ tịch HĐ Đạo đức]',
            'SO_QD_THANH_LAP_HD': '[số QĐ thành lập HĐ Đạo đức]',
            'NGAY_QD_HD': '[ngày]',
        }
        title = 'BẢN MẪU — QUYẾT ĐỊNH HỘI ĐỒNG ĐẠO ĐỨC (để điền)'
    else:
        F = {
            'TEN_BO_NGANH': 'BỘ GIÁO DỤC VÀ ĐÀO TẠO',
            'TEN_TRUONG': 'TRƯỜNG ĐẠI HỌC SƯ PHẠM HÀ NỘI',
            'TEN_HOI_DONG': 'Hội đồng Đạo đức trong Nghiên cứu Khoa học',
            'SO_QD': 'XXX/QĐ-ĐHSPHN',
            'DIA_DIEM': 'Hà Nội',
            'NGAY': '___', 'THANG': '___', 'NAM': '2025',
            'TEN_DE_TAI': '"Các yếu tố nguy cơ và bảo vệ đối với các rối loạn '
                          'lo âu ở học sinh trung học cơ sở Việt Nam"',
            'TEN_NCS': 'Công Thị Hằng',
            'DON_VI': 'Khoa Tâm lý Giáo dục, Trường Đại học Sư phạm Hà Nội',
            'HUONG_DAN': '[Học hàm/học vị] Nguyễn Minh Đức',
            'MUC_TIEU': 'Khảo sát mô hình tích hợp các yếu tố nguy cơ '
                        '(bị bắt nạt, áp lực học tập, nghiện điện thoại) và '
                        'các yếu tố bảo vệ (sự gắn bó với trường, sự hỗ trợ '
                        'từ cha mẹ, sự hỗ trợ từ bạn bè, lòng tự trọng) đối '
                        'với ba phân loại rối loạn lo âu (Lo âu Lan tỏa, '
                        'Lo âu Xã hội, Lo âu Chia ly) theo DSM-5 ở học sinh '
                        'trung học cơ sở Việt Nam',
            'DOI_TUONG': 'Học sinh trung học cơ sở (tuổi 11-14, lớp 6 đến '
                         'lớp 9)',
            'CO_MAU': '1.352 học sinh',
            'DIA_BAN': 'Trường Trung học cơ sở Nhật Tân và Trường Trung học '
                       'cơ sở Tây Mỗ, Hà Nội',
            'THOI_GIAN_BD': '___/2024',
            'THOI_GIAN_KT': '___/2025',
            'PHAM_VI_CONG_BO': 'Luận án tiến sĩ của nghiên cứu sinh Công Thị '
                               'Hằng và ba bài báo khoa học công bố trên tạp '
                               'chí quốc tế thuộc danh mục Scopus/WoS, sử '
                               'dụng cùng bộ dữ liệu thu thập theo đề cương '
                               'đã được phê duyệt: (i) bài Q2 về mô hình '
                               'phương trình cấu trúc tích hợp; (ii) bài Q3 '
                               'về bất biến cấu trúc theo giới; (iii) bài Q4 '
                               'về phân tích hồ sơ tiềm ẩn',
            'CHU_TICH_HD': '[Học hàm/học vị + Họ tên] Chủ tịch HĐ Đạo đức',
            'SO_QD_THANH_LAP_HD': 'XXX/QĐ-ĐHSPHN',
            'NGAY_QD_HD': '___',
        }
        title = 'BẢN ĐIỀN SẴN — QUYẾT ĐỊNH HỘI ĐỒNG ĐẠO ĐỨC (ví dụ cho NCS)'

    # Header — 2 cot
    table = d.add_table(rows=1, cols=2)
    table.autofit = False
    table.columns[0].width = Cm(7.5)
    table.columns[1].width = Cm(7.5)
    cell1 = table.cell(0, 0); cell2 = table.cell(0, 1)

    def add_centered_bold(cell, lines, sz=12):
        for ln in lines:
            p = cell.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_after = Pt(2)
            r = p.add_run(ln); r.font.name = 'Times New Roman'
            r.font.size = Pt(sz); r.bold = True

    # Clear default cell paragraph
    cell1.paragraphs[0]._element.getparent().remove(cell1.paragraphs[0]._element)
    cell2.paragraphs[0]._element.getparent().remove(cell2.paragraphs[0]._element)

    add_centered_bold(cell1, [F['TEN_BO_NGANH']], 11)
    add_centered_bold(cell1, [F['TEN_TRUONG']], 11)
    p = cell1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—————'); r.font.size = Pt(11)
    add_centered_bold(cell1, [F['TEN_HOI_DONG']], 11)
    p = cell1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'Số: {F["SO_QD"]}'); r.font.size = Pt(11); r.bold = True

    add_centered_bold(cell2, ['CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM'], 11)
    add_centered_bold(cell2, ['Độc lập - Tự do - Hạnh phúc'], 11)
    p = cell2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('—————'); r.font.size = Pt(11)
    p = cell2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'{F["DIA_DIEM"]}, ngày {F["NGAY"]} tháng {F["THANG"]} '
                  f'năm {F["NAM"]}')
    r.font.size = Pt(11); r.italic = True

    # Title QD
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run('QUYẾT ĐỊNH'); r.bold = True; r.font.size = Pt(15)

    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run('V/v: Phê duyệt về mặt khoa học và đạo đức đối với '
                  'đề cương nghiên cứu')
    r.italic = True; r.font.size = Pt(12)

    # Cancun cu
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'CHỦ TỊCH {F["TEN_HOI_DONG"].upper()}')
    r.bold = True; r.font.size = Pt(13)

    def CC(text):
        p = d.add_paragraph()
        p.paragraph_format.first_line_indent = Cm(1.0)
        p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(text); r.font.size = Pt(12); r.italic = True

    CC(f'Căn cứ Thông tư số 43/2024/TT-BYT ngày 12/12/2024 của Bộ '
       f'trưởng Bộ Y tế quy định việc thành lập, tổ chức và hoạt động '
       f'của Hội đồng đạo đức trong nghiên cứu y sinh học;')
    CC(f'Căn cứ Quyết định số {F["SO_QD_THANH_LAP_HD"]} ngày '
       f'{F["NGAY_QD_HD"]} của Hiệu trưởng {F["TEN_TRUONG"]} về việc '
       f'thành lập {F["TEN_HOI_DONG"]};')
    CC(f'Căn cứ Hồ sơ xin phê duyệt đạo đức đề tài nghiên cứu của '
       f'{F["TEN_NCS"]} đề ngày … tháng … năm …;')
    CC(f'Căn cứ kết luận của Hội đồng Đạo đức tại Phiên họp thẩm định '
       f'ngày {F["NGAY"]}/{F["THANG"]}/{F["NAM"]};')
    CC('Theo đề nghị của Thư ký Hội đồng Đạo đức,')

    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('QUYẾT ĐỊNH:'); r.bold = True; r.font.size = Pt(14)

    def DIEU(num, t):
        p = d.add_paragraph()
        p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.first_line_indent = Cm(0.5)
        r = p.add_run(f'Điều {num}. ')
        r.bold = True; r.font.size = Pt(12)
        r2 = p.add_run(t)
        r2.font.size = Pt(12)

    DIEU(1, f'Phê duyệt đề cương nghiên cứu có tiêu đề {F["TEN_DE_TAI"]} '
            f'của nghiên cứu sinh {F["TEN_NCS"]} ({F["DON_VI"]}), dưới '
            f'sự hướng dẫn của {F["HUONG_DAN"]}, về mặt khoa học và '
            f'đạo đức trong nghiên cứu.')

    DIEU(2, f'Phạm vi phê duyệt:')

    def SUB(label, val):
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r = p.add_run(f'• {label}: '); r.bold = True; r.font.size = Pt(12)
        r2 = p.add_run(val); r2.font.size = Pt(12)

    SUB('Mục tiêu nghiên cứu', F['MUC_TIEU'])
    SUB('Đối tượng tham gia', F['DOI_TUONG'])
    SUB('Cỡ mẫu', F['CO_MAU'])
    SUB('Địa bàn nghiên cứu', F['DIA_BAN'])
    SUB('Thời gian triển khai',
        f'từ {F["THOI_GIAN_BD"]} đến {F["THOI_GIAN_KT"]}')
    SUB('Phạm vi công bố kết quả', F['PHAM_VI_CONG_BO'])

    DIEU(3, f'Trách nhiệm của Chủ nhiệm đề tài/Nghiên cứu sinh:')

    def BULL(text, last=False):
        """Bullet đơn (không có label:value)"""
        p = d.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(2)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        ending = '.' if last else ';'
        r = p.add_run(f'• {text}{ending}')
        r.font.size = Pt(12)

    BULL('Thu thập sự đồng ý bằng văn bản (informed consent) của '
         'cha/mẹ/người giám hộ và sự đồng thuận bằng văn bản (assent) '
         'của học sinh trước khi tiến hành thu dữ liệu')
    BULL('Bảo mật thông tin cá nhân của đối tượng nghiên cứu theo '
         'quy định hiện hành')
    BULL('Báo cáo tiến độ thực hiện đề tài hàng năm cho Hội đồng Đạo '
         'đức')
    BULL('Thông báo kịp thời cho Hội đồng Đạo đức bất kỳ sự kiện bất '
         'lợi nào hoặc thay đổi đối với đề cương đã được phê duyệt')
    BULL('Lưu trữ hồ sơ nghiên cứu trong thời hạn tối thiểu 05 năm '
         'kể từ ngày hoàn thành nghiên cứu')
    BULL('Tuân thủ Tuyên bố Helsinki (1964 và các sửa đổi) và Quy chế '
         'thực hành lâm sàng tốt (GCP)', last=True)

    DIEU(4, f'Hiệu lực quyết định: Quyết định này có hiệu lực kể từ '
            f'ngày ký. Phạm vi phê duyệt áp dụng cho toàn bộ thời '
            f'gian triển khai nghiên cứu nêu tại Điều 2 và bao trùm '
            f'tất cả các công bố khoa học phát sinh từ đề tài này.')

    if is_blank:
        DIEU(5, f'Nơi nhận: Chủ nhiệm đề tài/NCS; Người hướng dẫn; '
                f'[Đơn vị quản lý chuyên môn của NCS]; [Phòng/Ban Quản '
                f'lý Khoa học Công nghệ]; Lưu Văn thư - HĐ Đạo đức.')
    else:
        DIEU(5, f'Nơi nhận: NCS Công Thị Hằng; Người hướng dẫn '
                f'(thầy Nguyễn Minh Đức); Khoa Tâm lý Giáo dục - '
                f'ĐHSPHN; Phòng Quản lý Khoa học Công nghệ - ĐHSPHN; '
                f'Lưu Văn thư - HĐ Đạo đức.')

    # Signature
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('CHỦ TỊCH HỘI ĐỒNG ĐẠO ĐỨC'); r.bold = True; r.font.size = Pt(12)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run('(Ký, ghi rõ họ tên và đóng dấu)'); r.italic = True; r.font.size = Pt(11)
    for _ in range(4):
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        p.paragraph_format.space_after = Pt(2)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r = p.add_run(F['CHU_TICH_HD']); r.bold = True; r.font.size = Pt(12)

    # Footer note
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    r = p.add_run(f'Ghi chú: ')
    r.bold = True; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
    if is_blank:
        r2 = p.add_run('Đây là BẢN MẪU để các cơ sở nghiên cứu tham '
                       'khảo. Các trường có dấu [...] cần được điền '
                       'theo thông tin thực tế của đề tài cụ thể. '
                       'Mẫu này được tổng hợp dựa trên Thông tư '
                       '43/2024/TT-BYT (hiệu lực 01/02/2025) và tham '
                       'khảo cấu trúc QĐ từ ĐH Y Hà Nội + BV Nguyễn '
                       'Tri Phương. Lưu ý: Thông tư 43/2024 chính '
                       'thức áp dụng cho nghiên cứu y sinh học do '
                       'Bộ Y tế quản lý; đối với nghiên cứu khoa '
                       'học giáo dục/tâm lý học do Bộ Giáo dục và '
                       'Đào tạo quản lý, các cơ sở có thể áp dụng '
                       'theo loại suy hoặc theo Quy chế nội bộ '
                       'riêng của trường — NCS cần xác minh với '
                       'Phòng Khoa học Công nghệ của trường mình.')
    else:
        r2 = p.add_run('Đây là BẢN ĐIỀN SẴN ví dụ minh họa cho đề '
                       'tài LA + 3 bài báo Q2/Q3/Q4 của NCS Công Thị '
                       'Hằng. Các trường còn dấu [...] hoặc XXX vẫn '
                       'cần được điền thông tin chính xác (số QĐ, '
                       'ngày, tên đầy đủ + học hàm học vị Chủ tịch '
                       'HĐ + Người hướng dẫn) sau khi nhận thư '
                       'chính thức từ Phòng KHCN HNUE.')
    r2.italic = True; r2.font.size = Pt(10)
    r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    # Clean metadata
    cp = d.core_properties
    cp.author = ''; cp.title = ''; cp.subject = ''
    cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
    cp.category = ''; cp.identifier = ''; cp.version = ''
    d.save(out_path)
    print(f'SAVED: {out_path}')


# ===== Build 2 ban =====
OUT_BLANK = os.path.join(ROOT, 'bai-bao-Q1',
                         'MauQD_HoiDongDaoDuc_BANMAU_08062026.docx')
OUT_FILLED = os.path.join(ROOT, 'bai-bao-Q1',
                          'MauQD_HoiDongDaoDuc_DIENSAN_08062026.docx')

build_qd(OUT_BLANK, is_blank=True)
build_qd(OUT_FILLED, is_blank=False)
