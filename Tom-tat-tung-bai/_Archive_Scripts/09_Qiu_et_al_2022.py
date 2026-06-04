# -*- coding: utf-8 -*-
"""Bài 9: Qiu et al. 2022 — CTH v5 (≤2,5 trang)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.5
for s in doc.sections: s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)

def rb(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

bl('Tóm tắt bài 9', bold=True)

rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình \u00ab Mối liên hệ giữa phong cách nuôi dạy con và khả năng phục hồi với triệu chứng trầm cảm và lo âu ở học sinh trung học cơ sở Trung Quốc \u00bb (Associations of Parenting Style and Resilience With Depression and Anxiety Symptoms in Chinese Middle School Students), do Zhihai Qiu, Ying Guo, Jun Wang và Hongbo Zhang (2022), khảo sát trên 2.079 học sinh (tuổi TB 16,7) tại thành phố Hợp Phì, tỉnh An Huy, Trung Quốc, xuất bản trên Frontiers in Psychology, 13:897339.')

rb('Phương pháp nghiên cứu')
bl('Công trình sử dụng Thang đo Trầm cảm CES-D (20 mục, ngưỡng ≥20, \u03b1 = 0,910) và Thang đo Lo âu Tự đánh giá SAS (20 mục, ngưỡng ≥50, \u03b1 = 0,799), kết hợp Thang đo Khả năng Phục hồi SRSMSS (\u03b1 = 0,931) và Thang đo Phong cách Nuôi dạy EMBU (9 chiều). Nói cách khác, nghiên cứu sử dụng phân tích hồ sơ tiềm ẩn (LPA) cho phong cách nuôi dạy kết hợp hồi quy logistic đa biến.', bold=True)
bl('Lấy mẫu cụm tháng 9\u201310/2021 tại Hợp Phì. Tỷ lệ phản hồi 98,06%. LPA xác định 3 hồ sơ nuôi dạy: Tích cực (58,6%), Trung bình (32,2%), Tiêu cực (9,1%).')

rb('Kết quả nghiên cứu định lượng')
tbl(
    ['Chỉ số', 'Tỷ lệ', 'Nam', 'Nữ', 'P'],
    [
        ['Trầm cảm (CES-D ≥20)', '26,0%', '24,3%', '28,9%', '<0,05'],
        ['Lo âu (SAS ≥50)', '13,4%', '11,1%', '17,5%', '<0,01'],
    ],
    widths=[4.0, 2.0, 2.0, 2.0, 2.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan (hồi quy logistic đa biến):', bold=True)
tbl(
    ['Yếu tố', 'OR trầm cảm', 'OR lo âu'],
    [
        ['Khả năng phục hồi thấp (vs cao)', '6,74', '2,80'],
        ['Nuôi dạy tiêu cực (vs trung bình)', '1,82', '2,01'],
        ['Nuôi dạy tích cực (vs trung bình)', '0,30 (bảo vệ)', '0,32 (bảo vệ)'],
    ],
    widths=[5.0, 3.5, 3.5]
)
doc.add_paragraph()

bh3('Khác biệt giới tính : Nữ có tỷ lệ lo âu (17,5%) cao hơn nam (11,1%, P < 0,01) \u2014 phù hợp xu hướng quốc tế.')

rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bl('*Khả năng phục hồi thấp* là yếu tố nguy cơ mạnh nhất cho trầm cảm (OR = 6,74) \u2014 gợi ý can thiệp tăng cường khả năng phục hồi (resilience) có thể hiệu quả. *Nuôi dạy tích cực* giảm nguy cơ trầm cảm 70% (OR = 0,30) và lo âu 68% (OR = 0,32), tương tự phát hiện của Pham và cộng sự (2024) tại Việt Nam (chăm sóc cảm xúc beta = \u20130,40).')

rb('Kết luận')
bl('Trầm cảm 26% và lo âu 13,4% ở học sinh Trung Quốc. Khả năng phục hồi thấp (OR = 6,74) và nuôi dạy tiêu cực (OR = 2,01) là yếu tố nguy cơ mạnh nhất. Nuôi dạy tích cực là yếu tố bảo vệ quan trọng. Can thiệp cần nhắm vào cả gia đình (cải thiện phong cách nuôi dạy) và cá nhân (tăng cường resilience).', bold=True)

rh2('Phản biện')
bl('Frontiers in Psychology (Q2, PubMed-indexed), LPA nâng cao. Tuy nhiên, cỡ mẫu chỉ ở 1 thành phố, nam chiếm 63,4% (thiên lệch giới), thiết kế cắt ngang, tương tác nuôi dạy × resilience không đạt ý nghĩa thống kê.')

rh2('Hướng nghiên cứu tiếp theo')
bl('Nghiên cứu dọc đánh giá tác động dài hạn của phong cách nuôi dạy. Can thiệp đào tạo phụ huynh kỹ năng nuôi dạy tích cực và đánh giá hiệu quả. Mở rộng đến nhiều thành phố với tỷ lệ giới cân bằng hơn.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50 Trung bình-Khá. LPA nâng cao, nhiều công cụ xác thực, nhưng 1 thành phố và thiên lệch giới.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('09_Qiu_et_al_2022.docx')
print('09_Qiu_et_al_2022.docx OK')
