# -*- coding: utf-8 -*-
"""Bai hoc song ngu 10 trang nhap mon tam ly hoc cho de tai
"Lo au vi thanh nien Viet Nam" — danh cho nguoi ngoai dao (thay NMD).
Soan 08/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'BaiHoc_NhapMonTLH_10trang_SongNgu_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def TITLE(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(18); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def PAGE_HEAD(num, en, vn):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(f'PAGE {num} — {en}'); r.font.name = 'Times New Roman'
    r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(f'TRANG {num} — {vn}'); r.font.name = 'Times New Roman'
    r.font.size = Pt(15); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def SUB(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def EN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def VN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def PB(t_en, t_vn):
    """Paragraph bilingual — EN then VN"""
    EN(t_en); VN(t_vn)

def BB(en, vn, level=0):
    """Bullet bilingual"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('• ' + en); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('▸ ' + vn); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def BOX(en, vn):
    """Boxed key concept"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(8)
    r = p.add_run('▣ KEY CONCEPT — ' + en); r.font.name = 'Times New Roman'
    r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('▣ KHÁI NIỆM CỐT LÕI — ' + vn); r.font.name = 'Times New Roman'
    r.font.size = Pt(11); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)

def PAGE_BREAK():
    d.add_page_break()


# ============================================================
TITLE('NHẬP MÔN TÂM LÝ HỌC CHO ĐỀ TÀI')
TITLE('"Lo âu Vị thành niên Việt Nam"')
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Introduction to Psychology for the Adolescent Anxiety Project')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True
r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Bilingual primer for non-specialist readers — 10 pages')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tài liệu nhập môn song ngữ Anh-Việt dành cho người ngoại đạo — 10 trang')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
# GLOSSARY page
PAGE_BREAK()
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TỪ ĐIỂN THUẬT NGỮ NHANH / QUICK GLOSSARY')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(10)
r = p.add_run('Cứ giở trang này khi gặp từ chưa rõ. Read this page first or whenever a term is unclear.')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
r.font.color.rgb = RGBColor(0x50, 0x50, 0x50)

glossary = [
    ('Anxiety vs Anxiety Disorder',
     'Lo âu vs Rối loạn lo âu',
     'Anxiety = a normal emotion (worry before exam). Anxiety disorder = persistent worry (≥6 months), excessive, and impairing daily life.',
     'Lo âu = cảm xúc bình thường (lo trước kỳ thi). Rối loạn lo âu = lo lắng dai dẳng (≥6 tháng), quá mức, làm suy giảm sinh hoạt hàng ngày.'),
    ('RCADS', 'Thang RCADS',
     'Revised Children\'s Anxiety and Depression Scale — a questionnaire designed for ages 8–18 that produces separate scores for GAD, SAD, SocAD and other subtypes.',
     'Thang đo lo âu và trầm cảm trẻ em phiên bản sửa đổi — bảng hỏi cho 8–18 tuổi, cho điểm riêng cho từng phân loại lo âu (GAD, SAD, SocAD và các loại khác).'),
    ('β (beta)', 'Hệ số β (beta)',
     'A standardized pathway coefficient from −1 to +1. Near 0 = weak effect; near ±1 = strong effect. β = 0.376 means a clearly meaningful effect; β = 0.05 would be trivial.',
     'Hệ số đường dẫn đã chuẩn hóa, có giá trị từ −1 tới +1. Gần 0 = hiệu ứng yếu; gần ±1 = hiệu ứng mạnh. β = 0,376 = hiệu ứng có ý nghĩa; β = 0,05 = không đáng kể.'),
    ('R² (R squared)', 'Hệ số R² (R bình phương)',
     'The proportion of variance explained. R² = 0.598 means the model explains 60% of why anxiety scores differ between students.',
     'Tỷ lệ phương sai được giải thích. R² = 0,598 nghĩa là mô hình giải thích 60% lý do điểm lo âu của học sinh khác nhau.'),
    ('SEM', 'SEM (mô hình phương trình cấu trúc)',
     'Structural Equation Modeling — a statistical method that tests many cause-effect pathways at once in a single model.',
     'Phương pháp thống kê kiểm tra đồng thời nhiều đường dẫn nguyên nhân-hậu quả trong một mô hình duy nhất.'),
    ('LPA', 'LPA (phân tích hồ sơ tiềm ẩn)',
     'Latent Profile Analysis — a method that groups individuals into "types" (e.g., high-stress vs low-stress students) based on patterns of measured variables.',
     'Phương pháp gộp các cá nhân thành các "loại" (ví dụ học sinh cao áp lực vs thấp áp lực) dựa trên mẫu hình các biến đã đo.'),
    ('DSM-5 / ICD-11', 'DSM-5 / ICD-11',
     'Two international classification systems for mental disorders. DSM-5 (2013, American Psychiatric Association) and ICD-11 (2022, World Health Organization).',
     'Hai hệ phân loại rối loạn tâm thần quốc tế. DSM-5 (2013, Hiệp hội Tâm thần học Hoa Kỳ) và ICD-11 (2022, Tổ chức Y tế Thế giới).'),
    ('Q1 / Q2 ranking', 'Q1 / Q2 ranking',
     'Quartile ranking of academic journals. Q1 = top 25% by impact; Q2 = next 25%. Higher Q = harder to publish in.',
     'Phân hạng theo tứ phân vị của tạp chí học thuật. Q1 = 25% đầu theo hệ số tác động; Q2 = 25% tiếp. Q càng cao = càng khó công bố.'),
    ('Validated scale', 'Thang đo đã chuẩn hóa',
     'A questionnaire whose reliability and validity have been tested. We use 8 such scales adapted to Vietnamese.',
     'Bảng hỏi đã được kiểm chứng độ tin cậy và tính giá trị. Chúng tôi sử dụng 8 thang đo loại này đã được chuyển ngữ sang tiếng Việt.'),
]
for term_en, term_vn, def_en, def_vn in glossary:
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(term_en + ' / ' + term_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(def_en)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(def_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)


# ===================== PAGE 1 =====================
PAGE_BREAK()
PAGE_HEAD(1, 'Why this topic matters', 'Tại sao đề tài này quan trọng')

SUB('1.1 — Global context / Bối cảnh toàn cầu')
PB('Mental disorders affect about 13% of children and adolescents '
   'worldwide, with anxiety disorders making up the largest share. '
   'Half of all lifetime mental disorders begin by age 14 (Kessler '
   '2005), and the global prevalence in young people has risen '
   'sharply since 2010 (Kieling 2024).',
   'Trên toàn thế giới, khoảng 13% trẻ em và vị thành niên mắc rối '
   'loạn tâm thần, trong đó nhóm rối loạn lo âu chiếm tỷ trọng lớn '
   'nhất. Một nửa số ca rối loạn tâm thần khởi phát trước tuổi 14 '
   '(Kessler 2005), và tỷ lệ mắc bệnh toàn cầu ở thanh thiếu niên '
   'đã gia tăng rõ rệt từ năm 2010 (Kieling 2024).')

SUB('1.2 — Vietnamese context / Bối cảnh Việt Nam')
PB('Vietnamese lower secondary students (ages 11–14) face a unique '
   'combination of pressures: high academic competition (Confucian '
   'culture), rapid digital connectivity, and tightening gender '
   'expectations during early adolescence. This is exactly the '
   'period when anxiety disorders are most likely to first appear.',
   'Học sinh trung học cơ sở Việt Nam (11–14 tuổi) đối diện với '
   'một tổ hợp áp lực đặc thù: cạnh tranh học tập cao theo văn hóa '
   'Nho giáo, kết nối số bùng nổ, và kỳ vọng giới ngày càng nghiêm '
   'ngặt ở đầu vị thành niên. Đây cũng chính là giai đoạn dễ khởi '
   'phát rối loạn lo âu nhất.')

SUB('1.3 — What this project does / Đề tài này làm gì')
PB('NCS Công Thị Hằng surveyed 1,352 lower secondary students in '
   'Hanoi (Nhat Tan and Tay Mo schools) using 8 validated scales. '
   'The data feed three planned international papers: Q2 (integrated '
   'model), Q3 (gender-specific pathways), and Q4 (student profile '
   'typology).',
   'NCS Công Thị Hằng khảo sát 1.352 học sinh THCS Hà Nội (trường '
   'Nhật Tân và Tây Mỗ) dùng 8 thang đo đã được chuẩn hóa. Bộ dữ '
   'liệu này sẽ phục vụ ba bài báo quốc tế đã lên kế hoạch: Q2 (mô '
   'hình tích hợp), Q3 (đường dẫn theo giới), Q4 (phân loại học sinh).')

BOX('Anxiety disorders are not "just stress" — they are clinically '
    'distinct conditions with specific diagnostic criteria, lasting '
    'months and impairing daily life.',
    'Rối loạn lo âu KHÔNG đơn thuần là "căng thẳng" — đó là tình '
    'trạng lâm sàng riêng biệt, có tiêu chí chẩn đoán cụ thể, kéo '
    'dài nhiều tháng và làm suy giảm chức năng sinh hoạt hàng ngày.')

SUB('1.4 — Meet our example student / Làm quen với học sinh ví dụ')
PB('Throughout this primer we use a hypothetical student, "Mai", to '
   'ground abstract ideas. Mai is 13 years old, in grade 7 at Tay '
   'Mo school in Hanoi. She studies hard, but recently has been '
   'avoiding class presentations and complaining of stomach aches '
   'on school mornings. Her parents wonder if it is just teenage '
   'mood or something more. Each page below will return to Mai '
   'briefly to show how the concept applies in real life.',
   'Xuyên suốt tài liệu này, chúng ta dùng một học sinh giả định '
   'tên "Mai" để làm rõ các khái niệm trừu tượng. Mai 13 tuổi, '
   'học lớp 7 trường Tây Mỗ Hà Nội. Em học chăm, nhưng gần đây hay '
   'tránh thuyết trình trước lớp và phàn nàn đau bụng vào sáng '
   'những ngày đi học. Cha mẹ Mai phân vân không biết đây là tâm '
   'trạng tuổi teen bình thường hay điều gì đó nghiêm trọng hơn. '
   'Mỗi trang dưới đây sẽ quay lại với Mai một chút để minh họa '
   'khái niệm áp dụng trong đời thực.')

SUB('1.5 — Self-check / Câu hỏi tự kiểm tra')
EN('Q: Why does adolescence (ages 11–18) deserve special attention '
   'in mental health research?')
VN('Hỏi: Tại sao tuổi vị thành niên (11–18) cần được chú ý đặc biệt '
   'trong nghiên cứu sức khỏe tâm thần?')
EN('A: Because half of all lifetime mental disorders begin by age '
   '14 (Kessler 2005), and the cohort is currently undergoing the '
   'fastest physical, social, and emotional development of life.')
VN('Đáp: Vì một nửa số ca rối loạn tâm thần khởi phát trước tuổi '
   '14 (Kessler 2005), và đây là giai đoạn phát triển nhanh nhất '
   'về thể chất, xã hội, và cảm xúc trong đời.')


# ===================== PAGE 2 =====================
PAGE_BREAK()
PAGE_HEAD(2, 'What is anxiety? Classification systems',
          'Lo âu là gì? Hệ phân loại quốc tế')

SUB('2.1 — Normal worry vs anxiety disorder / Lo lắng thường vs rối loạn lo âu')
PB('Everyone feels anxious sometimes — before an exam, a job '
   'interview, or a stage performance. This is normal and even '
   'helpful (it sharpens focus). An anxiety DISORDER is different: '
   'the worry is (1) persistent (lasting at least 6 months for '
   'most disorders), (2) disproportionate to actual threats, and '
   '(3) impairing — disrupting school, family, or social life.',
   'Ai cũng có lúc lo lắng — trước kỳ thi, phỏng vấn việc làm, '
   'hay biểu diễn trên sân khấu. Đây là lo lắng bình thường, '
   'thậm chí có ích (giúp tập trung tốt hơn). RỐI LOẠN lo âu thì '
   'khác: lo lắng (1) dai dẳng (thường tối thiểu 6 tháng), (2) '
   'không tương xứng với mức đe dọa thực tế, và (3) gây suy giảm '
   '— ảnh hưởng việc học, gia đình, hoặc đời sống xã hội.')

SUB('2.2 — Two main classification systems / Hai hệ phân loại chính')
PB('Worldwide, mental disorders are classified using two systems. '
   'DSM (Diagnostic and Statistical Manual) is published by the '
   'American Psychiatric Association — current edition is DSM-5-TR '
   '(2022), which contains the DSM-5 (2013, 237 disorders) plus 3 '
   'new diagnoses and revisions to over 70 criteria sets. ICD '
   '(International Classification of Diseases) is published by WHO '
   '— current edition ICD-11 Chapter 06 (effective worldwide from '
   '01/01/2022) contains 162 anxiety/stress-related categories.',
   'Trên thế giới, rối loạn tâm thần được phân loại theo hai hệ '
   'thống. DSM (Sổ tay Chẩn đoán và Thống kê các Rối loạn Tâm '
   'thần) do Hiệp hội Tâm thần học Hoa Kỳ phát hành — phiên bản '
   'hiện hành là DSM-5-TR (2022), gồm DSM-5 (2013, 237 rối loạn) '
   'cộng thêm 3 chẩn đoán mới và sửa hơn 70 bộ tiêu chuẩn (First '
   '2022). ICD (Phân loại Quốc tế các Bệnh) do WHO phát hành — '
   'phiên bản hiện hành ICD-11 Chương 06 (hiệu lực toàn cầu '
   'từ 01/01/2022) có 162 mục liên quan tới lo âu/stress (Pezzella '
   '2022).')

SUB('2.3 — Why this matters for our project / Tại sao quan trọng cho đề tài')
PB('NCS Công Thị Hằng uses DSM-5 categories to identify three '
   'anxiety subtypes (GAD, SAD, SocAD — see Page 3). Using a '
   'standard international system means the findings can be '
   'directly compared with research from other countries — a '
   'requirement for Q-ranked international journals.',
   'NCS Công Thị Hằng dùng phân loại DSM-5 để xác định ba phân '
   'loại lo âu (GAD, SAD, SocAD — xem Trang 3). Sử dụng hệ thống '
   'quốc tế chuẩn cho phép so sánh trực tiếp kết quả với nghiên '
   'cứu từ các quốc gia khác — yêu cầu bắt buộc của các tạp chí '
   'quốc tế thứ hạng Q.')


# ===================== PAGE 3 =====================
PAGE_BREAK()
PAGE_HEAD(3, 'Three anxiety subtypes in the project',
          'Ba phân loại lo âu trong đề tài')

SUB('3.1 — Generalized Anxiety Disorder (GAD) / Rối loạn lo âu lan tỏa')
PB('GAD is characterized by excessive worry about many different '
   'things — schoolwork, family health, the future — most days for '
   'at least 6 months. The worry feels uncontrollable and is '
   'accompanied by physical symptoms (muscle tension, sleep '
   'problems, fatigue). In our sample, NCS measured GAD using 7 '
   'items from the RCADS scale.',
   'GAD đặc trưng bởi lo lắng quá mức về nhiều thứ khác nhau — '
   'việc học, sức khỏe gia đình, tương lai — gần như hàng ngày '
   'trong tối thiểu 6 tháng. Lo lắng cảm thấy không kiểm soát '
   'được và đi kèm triệu chứng cơ thể (căng cơ, mất ngủ, mệt '
   'mỏi). Trong mẫu nghiên cứu, NCS đo GAD bằng 7 mục của thang '
   'RCADS.')

SUB('3.2 — Separation Anxiety Disorder (SAD) / Rối loạn lo âu chia ly')
PB('SAD is excessive fear of being apart from major attachment '
   'figures (usually parents). Adolescents with SAD worry that '
   'something bad will happen to a loved one, refuse to sleep '
   'away from home, or have nightmares about separation. Once '
   'thought to be "only a child disorder", SAD is now recognized '
   'as common and impairing in adolescents too. Measured with 4 '
   'items from RCADS.',
   'SAD là nỗi sợ quá mức khi xa cách người thân quan trọng '
   '(thường là cha mẹ). Vị thành niên mắc SAD lo lắng có chuyện '
   'xấu xảy ra với người thân, từ chối ngủ xa nhà, hay gặp ác '
   'mộng về sự chia ly. Trước đây cho là "rối loạn chỉ có ở trẻ '
   'em", SAD nay được công nhận xuất hiện phổ biến và gây suy '
   'giảm ở vị thành niên. Đo bằng 4 mục của RCADS.')

SUB('3.3 — Social Anxiety Disorder (SocAD) / Rối loạn lo âu xã hội')
PB('SocAD is intense fear of being judged or embarrassed in social '
   'situations — public speaking, meeting strangers, eating in '
   'public. Adolescents with SocAD may avoid school, refuse to '
   'speak in class, or become physically ill before social events. '
   'Lifetime prevalence reaches 12% in the US (Kessler 2005). '
   'Measured with 4 RCADS items.',
   'SocAD là nỗi sợ mãnh liệt bị đánh giá hay xấu hổ trong tình '
   'huống xã hội — thuyết trình, gặp người lạ, ăn nơi công cộng. '
   'Vị thành niên mắc SocAD có thể trốn học, từ chối phát biểu '
   'trong lớp, hoặc bị ốm thể chất trước sự kiện xã hội. Tỷ lệ '
   'mắc trong đời lên tới 12% ở Mỹ (Kessler 2005). Đo bằng 4 '
   'mục RCADS.')

BOX('All three subtypes share core anxiety machinery (excess fear '
    '+ avoidance) but differ in WHAT triggers fear. This distinction '
    'matters for intervention design.',
    'Cả ba phân loại đều chia sẻ cơ chế lo âu cốt lõi (sợ quá mức '
    '+ né tránh) nhưng KHÁC nhau ở yếu tố gây sợ. Sự khác biệt '
    'này quan trọng cho thiết kế can thiệp.')

SUB('3.4 — Back to Mai / Quay lại với Mai')
PB('Mai\'s avoidance of class presentations and physical symptoms '
   'before school suggest SocAD (Social Anxiety Disorder) is most '
   'likely — fear of being judged in social situations. She does '
   'not show separation anxiety (does not refuse to sleep alone) '
   'or generalized worry about many topics. A clinician would '
   'screen her using RCADS subscales to confirm.',
   'Việc Mai tránh thuyết trình và có triệu chứng cơ thể trước khi '
   'đến trường gợi ý nhiều khả năng là SocAD (Rối loạn Lo âu Xã '
   'hội) — sợ bị đánh giá trong tình huống xã hội. Mai không có '
   'biểu hiện lo âu chia ly (không từ chối ngủ một mình) hay lo '
   'lắng lan tỏa về nhiều chủ đề. Bác sĩ tâm lý sẽ sàng lọc bằng '
   'các tiểu thang RCADS để xác nhận.')

SUB('3.5 — Self-check / Câu hỏi tự kiểm tra')
EN('Q: How can a teacher tell SocAD apart from ordinary shyness?')
VN('Hỏi: Làm sao giáo viên có thể phân biệt SocAD với sự nhút '
   'nhát thông thường?')
EN('A: SocAD is persistent (months), excessive (fear far greater '
   'than the actual risk), and impairing (the student avoids '
   'school or fails academically). Shyness is mild, transient, '
   'and does not disrupt functioning.')
VN('Đáp: SocAD dai dẳng (nhiều tháng), quá mức (sợ vượt xa nguy '
   'cơ thực), và gây suy giảm (học sinh trốn học hoặc kết quả học '
   'tập kém). Nhút nhát thì nhẹ, thoáng qua, và không làm gián '
   'đoạn chức năng.')


# ===================== PAGE 4 =====================
PAGE_BREAK()
PAGE_HEAD(4, 'Three risk factors in our model',
          'Ba yếu tố nguy cơ trong mô hình')

SUB('4.1 — School bullying / Bắt nạt học đường')
PB('Bullying is repeated aggressive behavior aimed at causing harm '
   'to a peer with less power. It includes physical (hitting), '
   'verbal (insults, threats), and relational (exclusion, rumors) '
   'forms. NCS used the Olweus Bully/Victim Questionnaire (OBVQ, 8 '
   'items) to measure victimization. In our preliminary results, '
   'bullying showed the strongest pathway to Social Anxiety '
   'Disorder (β = 0.376).',
   'Bắt nạt là hành vi gây hại lặp đi lặp lại nhắm vào bạn cùng '
   'lứa có sức mạnh thấp hơn. Bao gồm bắt nạt thể chất (đánh đập), '
   'lời nói (xúc phạm, đe dọa), và quan hệ (cô lập, đồn đại). NCS '
   'dùng thang Olweus Bully/Victim (OBVQ, 8 mục) để đo nạn nhân '
   'hóa. Trong kết quả sơ bộ, bắt nạt cho thấy đường dẫn mạnh '
   'nhất tới Rối loạn Lo âu Xã hội (β = 0,376).')

SUB('4.2 — Academic stress / Áp lực học tập')
PB('Academic stress in Vietnamese adolescents is shaped by the '
   'Confucian "unforgiving culture" — high expectations from '
   'parents, teachers, and society about exam performance and '
   'university entry (Stankov 2010). NCS measured this using 4 '
   'items from the ESSA scale (Educational Stress Scale for '
   'Adolescents). Academic stress was the strongest predictor of '
   'total anxiety (β = 0.510).',
   'Áp lực học tập ở vị thành niên Việt Nam được định hình bởi '
   '"văn hóa Nho giáo nghiêm khắc" — kỳ vọng cao từ cha mẹ, giáo '
   'viên, và xã hội về thành tích kỳ thi và vào đại học (Stankov '
   '2010). NCS đo bằng 4 mục thang ESSA (Educational Stress Scale '
   'for Adolescents). Áp lực học tập là yếu tố dự báo mạnh nhất '
   'tới tổng RLLA (β = 0,510).')

SUB('4.3 — Smartphone addiction / Nghiện điện thoại thông minh')
PB('Smartphone addiction (also called "problematic smartphone use") '
   'is excessive use that disrupts sleep, study, or social '
   'relationships. Measured with 5 items from SAS-SV (Smartphone '
   'Addiction Scale-Short Version). This is the newest risk factor '
   'in adolescent mental health research, with mechanisms still '
   'being mapped — possibilities include sleep disruption, social '
   'comparison, and replacement of in-person interaction.',
   'Nghiện điện thoại thông minh (còn gọi "sử dụng điện thoại có '
   'vấn đề") là sử dụng quá mức gây gián đoạn giấc ngủ, học tập, '
   'hay quan hệ xã hội. Đo bằng 5 mục SAS-SV (Smartphone Addiction '
   'Scale-Short Version). Đây là yếu tố nguy cơ mới nhất trong '
   'nghiên cứu sức khỏe tâm thần vị thành niên, cơ chế đang được '
   'làm rõ — gồm gián đoạn giấc ngủ, so sánh xã hội, thay thế '
   'tương tác trực tiếp.')


# ===================== PAGE 5 =====================
PAGE_BREAK()
PAGE_HEAD(5, 'Four protective factors in our model',
          'Bốn yếu tố bảo vệ trong mô hình')

SUB('5.1 — School membership / Gắn bó với trường học')
PB('Psychological Sense of School Membership (PSSM, Goodenow 1993) '
   'is the feeling of being accepted, respected, and included by '
   'school peers and teachers. NCS used the 7-item PSSM scale. '
   'High school membership buffers against anxiety by providing '
   'a sense of safety and belonging.',
   'Cảm nhận Tâm lý về Gắn bó với Trường (PSSM, Goodenow 1993) là '
   'cảm giác được chấp nhận, tôn trọng, và bao gồm bởi bạn bè và '
   'giáo viên trong trường. NCS dùng thang PSSM 7 mục. Gắn bó '
   'trường cao có vai trò đệm chống lo âu bằng cách cung cấp cảm '
   'giác an toàn và thuộc về.')

SUB('5.2 — Parental support / Hỗ trợ từ cha mẹ')
PB('Parental support is the emotional and practical help adolescents '
   'receive from parents — listening, encouragement, financial '
   'support. Measured with 4 items from the MSPSS (Multidimensional '
   'Scale of Perceived Social Support). In Confucian families, '
   'parental support often takes the form of high expectations + '
   'sacrifice rather than emotional warmth.',
   'Hỗ trợ từ cha mẹ là sự giúp đỡ về cảm xúc và thực tế mà vị '
   'thành niên nhận được — lắng nghe, động viên, hỗ trợ tài chính. '
   'Đo bằng 4 mục MSPSS (Multidimensional Scale of Perceived '
   'Social Support). Trong gia đình Nho giáo, hỗ trợ cha mẹ thường '
   'biểu hiện qua kỳ vọng cao + hy sinh thay vì sự ấm áp cảm xúc '
   'trực tiếp.')

SUB('5.3 — Peer support / Hỗ trợ từ bạn bè')
PB('Peer support is help from friends — sharing problems, having '
   'fun together, mutual encouragement. Measured with 4 items '
   'from MSPSS. Counter-intuitively, in our Vietnamese sample, '
   'peer support showed almost no protective effect against '
   'anxiety. We propose two explanations: cultural restraint in '
   'sharing (tam giao tradition; Small & Blanc 2021) and co-'
   'rumination (Rose 2002).',
   'Hỗ trợ từ bạn bè là sự giúp đỡ từ bạn — chia sẻ vấn đề, vui '
   'chơi cùng nhau, động viên lẫn nhau. Đo bằng 4 mục MSPSS. Đáng '
   'ngạc nhiên là trong mẫu Việt Nam của chúng tôi, hỗ trợ bạn bè '
   'gần như không có hiệu ứng bảo vệ chống lo âu. Chúng tôi đề '
   'xuất hai cách giải thích: sự chừng mực văn hóa trong chia sẻ '
   '(truyền thống tam giáo; Small & Blanc 2021) và co-rumination '
   '(Rose 2002).')

SUB('5.4 — Self-esteem / Tự trọng')
PB('Self-esteem is the overall evaluation of one\'s own worth. '
   'Measured with 5 items from the Rosenberg Self-Esteem Scale '
   '(RSES). In our data, self-esteem was the strongest single '
   'protective factor (β = −0.455 against GAD). This makes self-'
   'esteem a high-yield target for prevention programs.',
   'Tự trọng là sự đánh giá tổng thể về giá trị bản thân. Đo bằng '
   '5 mục thang Rosenberg Self-Esteem (RSES). Trong dữ liệu, tự '
   'trọng là yếu tố bảo vệ đơn mạnh nhất (β = −0,455 chống GAD). '
   'Điều này khiến tự trọng trở thành mục tiêu can thiệp có hiệu '
   'suất cao.')


# ===================== PAGE 6 =====================
PAGE_BREAK()
PAGE_HEAD(6, 'Vietnamese cultural framework',
          'Khung lý thuyết văn hóa Việt Nam')

SUB('6.1 — Tam giao / Three Teachings / Tam giáo')
PB('Vietnamese culture is shaped by the coexistence of three '
   'philosophical-religious traditions: Confucianism (social '
   'hierarchy, education), Buddhism (compassion, endurance), '
   'and Taoism (harmony with nature). Small and Blanc (2021) '
   'called this "cultural additivity" — three traditions reinforce '
   'each other rather than compete. One shared value is "nhẫn" '
   '(emotional restraint, patience), which may reduce open '
   'sharing of distress.',
   'Văn hóa Việt Nam được định hình bởi sự cộng tồn của ba truyền '
   'thống triết học-tôn giáo: Nho giáo (thứ bậc xã hội, giáo dục), '
   'Phật giáo (lòng từ bi, nhẫn nhịn), và Đạo giáo (hòa hợp với '
   'thiên nhiên). Small và Blanc (2021) gọi đây là "cộng dồn văn '
   'hóa" — ba truyền thống củng cố lẫn nhau thay vì cạnh tranh. '
   'Một giá trị chung là "nhẫn" — sự chừng mực cảm xúc, kiên '
   'nhẫn, có thể làm giảm chia sẻ khó khăn một cách công khai.')

SUB('6.2 — Confucian academic culture / Văn hóa học thuật Nho giáo')
PB('Stankov (2010) described the Confucian academic culture as '
   '"unforgiving" — students from Confucian East Asian countries '
   '(China, Korea, Japan, Vietnam) achieve high academic '
   'performance but also experience higher test anxiety and self-'
   'doubt compared to Western peers. The pressure is dual: prove '
   'one\'s effort + maintain family honor.',
   'Stankov (2010) mô tả văn hóa học thuật Nho giáo là "khắc '
   'nghiệt" — học sinh từ các nước Đông Á Nho giáo (Trung Quốc, '
   'Hàn Quốc, Nhật Bản, Việt Nam) đạt thành tích học tập cao '
   'nhưng cũng trải nghiệm lo âu kiểm tra và tự nghi ngờ cao hơn '
   'so với bạn phương Tây. Áp lực kép: chứng minh sự cố gắng + '
   'giữ gìn thanh danh gia đình.')

SUB('6.3 — Co-rumination in girls / Co-rumination ở nữ sinh')
PB('Rose (2002) discovered that girls in early adolescence (grades '
   '3, 5, 7, 9; N = 608) engage in "co-rumination" — repeatedly '
   'discussing problems with friends, speculating about negative '
   'outcomes, and dwelling on bad feelings. This deepens friendship '
   'but also amplifies anxiety. Girls show this pattern more than '
   'boys, especially around ages 11–14.',
   'Rose (2002) phát hiện nữ sinh đầu vị thành niên (lớp 3, 5, '
   '7, 9; N = 608) tham gia "co-rumination" — thảo luận lặp đi '
   'lặp lại các vấn đề với bạn, suy đoán về kết quả tiêu cực, và '
   'đắm chìm trong cảm xúc xấu. Điều này làm sâu sắc tình bạn '
   'nhưng cũng khuếch đại lo âu. Nữ thể hiện kiểu này nhiều hơn '
   'nam, đặc biệt ở tuổi 11–14.')


# ===================== PAGE 7 =====================
PAGE_BREAK()
PAGE_HEAD(7, 'Structural Equation Modeling (SEM) — plain explanation',
          'Mô hình Phương trình Cấu trúc (SEM) — giải thích đơn giản')

SUB('7.1 — What SEM tests / SEM kiểm tra gì')
PB('Regular regression tests one relationship at a time: "Does '
   'bullying predict anxiety?" SEM tests many relationships at '
   'once: "Do bullying + academic stress + smartphone addiction '
   'jointly predict GAD + SAD + SocAD, while school membership '
   '+ parental support + peer support + self-esteem buffer '
   'against all three subtypes — all in one model?"',
   'Hồi quy thông thường kiểm tra một mối quan hệ tại một thời '
   'điểm: "Bắt nạt có dự báo lo âu không?" SEM kiểm tra nhiều '
   'mối quan hệ cùng lúc: "Bắt nạt + áp lực học tập + nghiện '
   'điện thoại có cùng dự báo GAD + SAD + SocAD, trong khi gắn '
   'bó trường + hỗ trợ cha mẹ + hỗ trợ bạn bè + tự trọng có vai '
   'trò đệm chống cả ba phân loại — tất cả trong cùng một mô '
   'hình không?"')

SUB('7.2 — Latent variables / Biến tiềm ẩn')
PB('A KEY idea in SEM is the "latent variable" — a concept we '
   'cannot measure directly (like "anxiety"), but only through '
   'observable items (questions on a scale). SEM lets us model '
   'the underlying concept separately from measurement error. '
   'For example, "GAD" is a latent variable measured by 7 RCADS '
   'items.',
   'Một ý tưởng CHỦ CHỐT trong SEM là "biến tiềm ẩn" — khái niệm '
   'không đo trực tiếp được (như "lo âu"), chỉ thông qua các mục '
   'quan sát được (câu hỏi trên thang đo). SEM cho phép mô hình '
   'hóa khái niệm gốc tách bạch khỏi sai số đo lường. Ví dụ '
   '"GAD" là biến tiềm ẩn đo bằng 7 mục RCADS.')

SUB('7.3 — Goodness-of-fit / Độ phù hợp mô hình')
PB('After building an SEM, we test whether it fits the data using '
   'standard indices: CFI ≥ 0.90 (Comparative Fit Index), RMSEA '
   '≤ 0.08 (Root Mean Square Error of Approximation), SRMR ≤ '
   '0.08 (Standardized Root Mean Square Residual). Higher CFI is '
   'better; lower RMSEA/SRMR is better. Thresholds come from '
   'Hu and Bentler (1999).',
   'Sau khi xây SEM, chúng ta kiểm tra mô hình có phù hợp dữ '
   'liệu hay không thông qua các chỉ số chuẩn: CFI ≥ 0,90 '
   '(Comparative Fit Index), RMSEA ≤ 0,08 (Root Mean Square '
   'Error of Approximation), SRMR ≤ 0,08 (Standardized Root '
   'Mean Square Residual). CFI cao là tốt; RMSEA/SRMR thấp là '
   'tốt. Ngưỡng từ Hu và Bentler (1999).')

BOX('SEM is the gold-standard method for testing complex multi-'
    'variable theories — it powers most modern Q1/Q2 papers in '
    'developmental psychology and psychiatry.',
    'SEM là phương pháp chuẩn vàng để kiểm chứng các lý thuyết đa '
    'biến phức tạp — đa số các bài Q1/Q2 hiện đại trong tâm lý '
    'học phát triển và tâm thần học đều dùng SEM.')

SUB('7.4 — Helpful analogy / Phép ẩn dụ giúp dễ hình dung')
PB('Think of SEM as a map of a city. Each road on the map is a '
   'pathway (β coefficient) between two locations (variables). '
   'Ordinary regression draws one road at a time. SEM draws the '
   'whole road network at once — showing not just whether each '
   'road exists, but whether the whole transport system makes '
   'sense as a unit. Goodness-of-fit indices are like asking: '
   'does our map match the real city? CFI ≥ 0.90 = a good '
   'enough match.',
   'Hãy tưởng tượng SEM như tấm bản đồ một thành phố. Mỗi con '
   'đường trên bản đồ là một đường dẫn (hệ số β) giữa hai địa '
   'điểm (biến). Hồi quy thông thường vẽ một con đường tại một '
   'thời điểm. SEM vẽ toàn bộ mạng lưới đường cùng lúc — cho '
   'biết không chỉ từng đường có tồn tại không, mà cả hệ thống '
   'giao thông có hợp lý không. Các chỉ số độ phù hợp giống như '
   'hỏi: bản đồ của chúng ta có khớp với thành phố thực không? '
   'CFI ≥ 0,90 = khớp đủ tốt.')

SUB('7.5 — Self-check / Câu hỏi tự kiểm tra')
EN('Q: What makes SEM more powerful than running 10 separate '
   'regressions?')
VN('Hỏi: Điều gì khiến SEM mạnh hơn việc chạy 10 hồi quy riêng lẻ?')
EN('A: SEM tests all relationships simultaneously so error '
   'compounds less, models measurement error in latent variables, '
   'and provides a single fit statistic showing whether the '
   'whole theory holds.',
   )
VN('Đáp: SEM kiểm tra tất cả mối quan hệ cùng lúc nên sai số ít '
   'tích lũy hơn, mô hình hóa được sai số đo lường trong biến '
   'tiềm ẩn, và cho một chỉ số phù hợp duy nhất cho biết toàn bộ '
   'lý thuyết có đúng không.')


# ===================== PAGE 8 =====================
PAGE_BREAK()
PAGE_HEAD(8, 'Q2 — Integrated higher-order SEM',
          'Q2 — SEM tích hợp bậc cao')

SUB('8.1 — Hypotheses / Giả thuyết')
PB('Q2 tests three main hypotheses. H1: All three risk factors '
   '(bullying, academic stress, smartphone addiction) have '
   'positive pathways to GAD, SAD, and SocAD. H2: All four '
   'protective factors have negative pathways. H3: The integrated '
   'higher-order model (Risk → Total RLLA + Protective → Total '
   'RLLA simultaneously) explains substantial variance (R² > '
   '0.40).',
   'Q2 kiểm chứng ba giả thuyết chính. H1: Cả ba yếu tố nguy cơ '
   '(bắt nạt, áp lực học tập, nghiện điện thoại) có đường dẫn '
   'dương tới GAD, SAD, SocAD. H2: Cả bốn yếu tố bảo vệ có đường '
   'dẫn âm. H3: Mô hình bậc cao tích hợp (Nguy cơ → Tổng RLLA + '
   'Bảo vệ → Tổng RLLA đồng thời) giải thích lượng phương sai '
   'lớn (R² > 0,40).')

SUB('8.2 — Key findings / Phát hiện chính')
PB('All three hypotheses were confirmed. Strongest risk pathway: '
   'bullying → SAD (β = 0.376). Strongest protective pathway: '
   'self-esteem → GAD (β = −0.455). Total integrated model: R² = '
   '0.598 — meaning the model explains about 60% of variance in '
   'total anxiety symptoms. The peer support pathway was '
   'unexpectedly weak.',
   'Cả ba giả thuyết được khẳng định. Đường dẫn nguy cơ mạnh '
   'nhất: bắt nạt → SAD (β = 0,376). Đường dẫn bảo vệ mạnh nhất: '
   'tự trọng → GAD (β = −0,455). Mô hình tích hợp tổng: R² = '
   '0,598 — nghĩa là mô hình giải thích khoảng 60% phương sai '
   'của tổng triệu chứng lo âu. Đường dẫn từ hỗ trợ bạn bè yếu '
   'một cách bất ngờ.')

SUB('8.3 — Why this matters / Vì sao quan trọng')
PB('Q2 is the first paper to integrate three risk + four '
   'protective factors → three DSM-5 anxiety subtypes in a single '
   'higher-order SEM for Vietnamese adolescents. It identifies '
   'self-esteem as a high-leverage intervention target. The '
   'paper is now planned for Frontiers in Psychiatry.',
   'Q2 là bài đầu tiên tích hợp ba yếu tố nguy cơ + bốn yếu tố '
   'bảo vệ → ba phân loại lo âu DSM-5 trong một SEM bậc cao cho '
   'vị thành niên Việt Nam. Bài xác định tự trọng là mục tiêu '
   'can thiệp đòn bẩy cao. Bài hiện được lên kế hoạch cho '
   'Frontiers in Psychiatry.')


# ===================== PAGE 9 =====================
PAGE_BREAK()
PAGE_HEAD(9, 'Q3 — Multi-group SEM by gender',
          'Q3 — SEM đa nhóm theo giới')

SUB('9.1 — Why a separate paper / Vì sao tách bài riêng')
PB('In our preliminary results, girls scored higher on Social '
   'Anxiety Disorder than boys. This was surprising given that '
   'Jefferies and Ungar (2020) — a large 7-country study of '
   '6,825 young adults aged 16–29 — found NO gender difference. '
   'Q3 will test whether the entire integrated SEM works the '
   'same way for boys and girls — that is, whether each pathway '
   '(β value) is identical across gender.',
   'Trong kết quả sơ bộ, nữ có điểm SAD cao hơn nam. Điều này '
   'gây ngạc nhiên vì Jefferies và Ungar (2020) — nghiên cứu '
   'lớn 7 quốc gia trên 6.825 thanh niên 16–29 tuổi — KHÔNG '
   'phát hiện khác biệt giới. Q3 sẽ kiểm chứng xem toàn bộ SEM '
   'tích hợp có hoạt động giống nhau cho nam và nữ — tức là '
   'từng đường dẫn (giá trị β) có giống nhau theo giới hay '
   'không.')

SUB('9.2 — Three levels of invariance / Ba mức bất biến')
PB('Q3 tests three sequential levels: (1) configural invariance '
   '— same factor structure in both groups; (2) metric '
   'invariance — same factor loadings; (3) scalar invariance '
   '— same intercepts. Failure at any level reveals where '
   'gender matters. The decision criterion is ΔCFI ≤ 0.01 '
   '(Cheung and Rensvold 2002).',
   'Q3 kiểm chứng ba mức tuần tự: (1) bất biến cấu hình '
   '(configural) — cùng cấu trúc nhân tố ở hai nhóm; (2) bất '
   'biến metric — cùng hệ số tải; (3) bất biến scalar — cùng '
   'hệ số chặn. Thất bại ở bất kỳ mức nào đều chỉ ra nơi giới '
   'có vai trò. Tiêu chí quyết định ΔCFI ≤ 0,01 (Cheung và '
   'Rensvold 2002).')

SUB('9.3 — Invariance analogy / Phép ẩn dụ về bất biến')
PB('Imagine two versions of the same city map — one for boys, one '
   'for girls. Configural invariance = both versions have the same '
   'kind of map (roads, intersections). Metric invariance = the '
   'roads have the same length in both versions. Scalar invariance '
   '= the starting point is at the same address. If any level '
   'fails, the maps differ in a specific way — and that is where '
   'gender matters.',
   'Hãy tưởng tượng hai phiên bản của cùng một bản đồ thành phố — '
   'một cho nam, một cho nữ. Bất biến cấu hình = cả hai phiên bản '
   'có cùng kiểu bản đồ (đường, giao lộ). Bất biến metric = các '
   'con đường có cùng độ dài ở cả hai phiên bản. Bất biến scalar '
   '= điểm xuất phát ở cùng địa chỉ. Nếu bất kỳ mức nào không đạt, '
   'hai bản đồ khác nhau theo cách cụ thể — và đó là nơi giới '
   'có vai trò.')

SUB('9.4 — Cultural-developmental framing / Khung phát triển-văn hóa')
PB('Q3 will provide six arguments for the gender difference '
   'observed in our sample but absent in Jefferies\': (1) the '
   '11–14 developmental window is a peak vulnerability period '
   'for girls (Hankin 2007; McLean and Anderson 2009); (2) co-'
   'rumination peaks in middle school girls (Rose 2002); (3) '
   'Confucian gender role pressure is strictest during school '
   'years (Stankov 2010; Small and Blanc 2021); (4) school-'
   'context measurement amplifies girls\' SAD; (5) Jefferies\' '
   'market-research panel has selection bias; (6) Jefferies\' '
   'age effect dominates and masks gender.',
   'Q3 sẽ đưa ra sáu luận chứng cho khác biệt giới quan sát '
   'trong mẫu nhưng vắng trong Jefferies: (1) cửa sổ phát '
   'triển 11–14 là đỉnh nguy cơ ở nữ (Hankin 2007; McLean và '
   'Anderson 2009); (2) co-rumination đỉnh ở nữ sinh THCS '
   '(Rose 2002); (3) áp lực vai trò giới Nho giáo nghiêm ngặt '
   'nhất ở tuổi học sinh (Stankov 2010; Small và Blanc 2021); '
   '(4) bối cảnh đo lường học đường khuếch đại SAD nữ; (5) '
   'mẫu market research Jefferies có sai lệch lựa chọn; (6) '
   'hiệu ứng tuổi trong Jefferies áp đảo và che giấu giới.')


# ===================== PAGE 10 =====================
PAGE_BREAK()
PAGE_HEAD(10, 'Q4 — Latent Profile Analysis (LPA)',
          'Q4 — Phân tích Hồ sơ Tiềm ẩn')

SUB('10.1 — Different question / Câu hỏi khác')
PB('Q2 and Q3 ask "what predicts anxiety?" Q4 asks a different '
   'question: "Are there distinct TYPES of students in our '
   'sample — each with a characteristic pattern of risk and '
   'protective factors?" For example, one type might be high-'
   'stress + low-support; another might be moderate-stress + '
   'high-self-esteem. These types are not visible from regression '
   'or SEM — they require a different statistical method.',
   'Q2 và Q3 hỏi "yếu tố nào dự báo lo âu?" Q4 hỏi câu khác: '
   '"Có các LOẠI học sinh riêng biệt trong mẫu — mỗi loại có '
   'mẫu hình đặc trưng về yếu tố nguy cơ và bảo vệ không?" Ví '
   'dụ, một loại có thể là cao áp lực + thấp hỗ trợ; loại khác '
   'có thể là vừa áp lực + cao tự trọng. Các loại này không '
   'lộ ra từ hồi quy hay SEM — chúng cần phương pháp thống '
   'kê khác.')

SUB('10.2 — How LPA works / LPA hoạt động ra sao')
PB('LPA (Latent Profile Analysis) is a person-centered method '
   'that groups individuals into unobserved profiles based on '
   'their pattern across many measured variables. It typically '
   'identifies 3–5 profiles that best fit the data. Each profile '
   'is interpreted by the relative levels of input variables. '
   'LPA is now widely used in adolescent mental health research.',
   'LPA (Latent Profile Analysis) là phương pháp lấy con người '
   'làm trung tâm, gộp các cá nhân thành các hồ sơ không quan '
   'sát được dựa trên mẫu hình của họ qua nhiều biến đo lường. '
   'Thường xác định 3–5 hồ sơ phù hợp nhất dữ liệu. Mỗi hồ sơ '
   'được diễn giải qua mức tương đối của các biến đầu vào. LPA '
   'hiện được sử dụng rộng rãi trong nghiên cứu sức khỏe tâm '
   'thần vị thành niên.')

SUB('10.3 — Why Q4 doesn\'t overlap Q2/Q3 / Vì sao Q4 không trùng Q2/Q3')
PB('Q2 + Q3 use SEM (variable-centered): each variable predicts '
   'each outcome. Q4 uses LPA (person-centered): each student '
   'belongs to a profile. The two approaches answer different '
   'scientific questions and use different statistical mechanics. '
   'There is no double-publication risk. Q4 fits journals such '
   'as Journal of Psychopathology and Behavioral Assessment or '
   'BMC Psychiatry. Q4 is planned for next year if the dataset '
   'still permits.',
   'Q2 + Q3 dùng SEM (lấy biến làm trung tâm): mỗi biến dự báo '
   'mỗi outcome. Q4 dùng LPA (lấy con người làm trung tâm): mỗi '
   'học sinh thuộc về một hồ sơ. Hai cách tiếp cận trả lời câu '
   'hỏi khoa học khác nhau và dùng cơ chế thống kê khác nhau. '
   'Không có rủi ro xuất bản trùng lặp. Q4 phù hợp các tạp chí '
   'như Journal of Psychopathology and Behavioral Assessment '
   'hoặc BMC Psychiatry. Q4 dự kiến năm sau nếu dữ liệu vẫn cho '
   'phép.')

BOX('Three papers from one large dataset, each with a different '
    'angle: Q2 (integrated mechanisms), Q3 (gender invariance), '
    'Q4 (student typology). This is a "hat-trick" pattern common '
    'in high-impact cohort programs.',
    'Ba bài từ một bộ dữ liệu lớn, mỗi bài một góc độ khác '
    'nhau: Q2 (cơ chế tích hợp), Q3 (bất biến giới), Q4 (phân '
    'loại học sinh). Đây là mẫu hình "hat-trick" phổ biến '
    'trong các chương trình cohort có tầm ảnh hưởng cao.')

SUB('10.4 — LPA analogy / Phép ẩn dụ về LPA')
PB('LPA is like sorting a library. SEM asks "does this book\'s '
   'thickness predict its number of pages?" — a variable-by-variable '
   'question. LPA asks "are these books naturally grouped into '
   'shelves — fiction, science, history — based on patterns of '
   'multiple features together?" In our data, LPA may reveal '
   'student profiles such as "high-stress + low-support + high-'
   'anxiety" vs "moderate-stress + high-self-esteem + low-anxiety".',
   'LPA giống như sắp xếp thư viện. SEM hỏi "độ dày của cuốn sách '
   'có dự báo số trang của nó không?" — câu hỏi từng biến một. '
   'LPA hỏi "các cuốn sách có tự nhiên phân nhóm thành các kệ — '
   'tiểu thuyết, khoa học, lịch sử — dựa trên mẫu hình của nhiều '
   'đặc điểm cùng lúc không?" Trong dữ liệu, LPA có thể phát hiện '
   'các hồ sơ học sinh như "cao áp lực + thấp hỗ trợ + cao lo âu" '
   'vs "vừa áp lực + cao tự trọng + thấp lo âu".')

SUB('10.5 — Self-check / Câu hỏi tự kiểm tra')
EN('Q: Why doesn\'t LPA double-publish on Q2/Q3 data?')
VN('Hỏi: Vì sao LPA không bị tính là tự đạo văn từ dữ liệu Q2/Q3?')
EN('A: Variable-centered (SEM) and person-centered (LPA) answer '
   'fundamentally different scientific questions and use different '
   'statistical mechanics — like asking "how tall is the average '
   'tree?" vs "what kinds of trees grow here?"')
VN('Đáp: Cách tiếp cận lấy biến làm trung tâm (SEM) và lấy con '
   'người làm trung tâm (LPA) trả lời các câu hỏi khoa học khác '
   'nhau về bản chất và dùng cơ chế thống kê khác nhau — giống '
   'như hỏi "cây trung bình cao bao nhiêu?" vs "có những loài cây '
   'gì mọc ở đây?"')


# ============================================================
PAGE_BREAK()
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('KEY REFERENCES / THAM KHẢO CHÍNH')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

refs = [
    '[1] Kieling C, Buchweitz C, Caye A, et al. (2024). Worldwide '
    'prevalence and disability from mental disorders across '
    'childhood and adolescence. JAMA Psychiatry 81(4):347-356. '
    'DOI: 10.1001/jamapsychiatry.2023.5051.',
    '[2] Kessler RC, Berglund P, Demler O, et al. (2005). Lifetime '
    'prevalence and age-of-onset distributions of DSM-IV disorders '
    'in the NCS-R. Arch Gen Psychiatry 62(6):593-602. '
    'PMID: 15939837.',
    '[3] First MB, Yousif LH, Clarke DE, Wang PS, Gogtay N, '
    'Appelbaum PS. (2022). DSM-5-TR: overview of what\'s new and '
    'what\'s changed. World Psychiatry 21(2):218-219. '
    'PMID: 35524596.',
    '[4] Pezzella P. (2022). The ICD-11 is now officially in '
    'effect. World Psychiatry 21(2):331-332. PMID: 35524598.',
    '[5] Hankin BL, Mermelstein R, Roesch L. (2007). Sex '
    'differences in adolescent depression: Stress exposure and '
    'reactivity models. Child Development 78(1):279-295. '
    'PMID: 17328705.',
    '[6] McLean CP, Anderson ER. (2009). Brave men and timid '
    'women? A review of the gender differences in fear and '
    'anxiety. Clinical Psychology Review 29(6):496-505. '
    'PMID: 19541399.',
    '[7] Rose AJ. (2002). Co-rumination in the friendships of '
    'girls and boys. Child Development 73(6):1830-1843. '
    'PMID: 12487497.',
    '[8] Stankov L. (2010). Unforgiving Confucian culture: A '
    'breeding ground for high academic achievement, test anxiety '
    'and self-doubt? Learning and Individual Differences '
    '20(6):555-563.',
    '[9] Small S, Blanc J. (2021). Mental Health During COVID-19: '
    'Tam Giao and Vietnam\'s Response. Frontiers in Psychiatry '
    '11:589618. PMID: 33536961.',
    '[10] Jefferies P, Ungar M. (2020). Social anxiety in young '
    'people: A prevalence study in seven countries. PLOS ONE '
    '15(9):e0239133. DOI: 10.1371/journal.pone.0239133.',
    '[11] Hu LT, Bentler PM. (1999). Cutoff criteria for fit '
    'indexes in covariance structure analysis. Structural '
    'Equation Modeling 6(1):1-55.',
    '[12] Cheung GW, Rensvold RB. (2002). Evaluating goodness-of-'
    'fit indexes for testing measurement invariance. Structural '
    'Equation Modeling 9(2):233-255.',
]
for r in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    rr = p.add_run(r); rr.font.name = 'Times New Roman'; rr.font.size = Pt(10)

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Soạn 08/06/2026 — Bài học nhập môn 10 trang')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')

from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            chunks.append(cell.text)
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
