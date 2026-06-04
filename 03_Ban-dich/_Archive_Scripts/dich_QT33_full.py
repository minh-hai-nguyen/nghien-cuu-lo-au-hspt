# -*- coding: utf-8 -*-
"""Dịch đầy đủ QT33 — JAMA 2024 — Schmidt-Persson et al. — Screen Media RCT"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link bài báo gốc: https://doi.org/10.1001/jamanetworkopen.2024.19881', size=10)

add_heading(doc, 'Sử dụng phương tiện màn hình và sức khỏe tâm thần của trẻ em và thanh thiếu niên: Phân tích thứ cấp thử nghiệm lâm sàng ngẫu nhiên', 1)
h = doc.add_paragraph()
r = h.add_run('Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial'),
    ('Tiêu đề dịch', 'Sử dụng phương tiện màn hình và sức khỏe tâm thần của trẻ em và thanh thiếu niên: Phân tích thứ cấp thử nghiệm lâm sàng ngẫu nhiên'),
    ('Tác giả', 'Jesper Schmidt-Persson, PhD; Martin Gillies Banke Rasmussen, PhD; Sarah Overgaard Sørensen, MSc; Sofie Rath Mortensen, MSc; Line Grønholt Olesen, PhD; Søren Brage, PhD; Peter Lund Kristensen, PhD; Niels Bilenberg, PhD; Anders Grøntved, PhD'),
    ('Cơ quan', 'Research Unit for Exercise Epidemiology, Đại học Southern Denmark; MRC Epidemiology Unit, Đại học Cambridge; Department of Child and Adolescent Psychiatry, Đại học Southern Denmark'),
    ('Tạp chí', 'JAMA Network Open (Q1, IF ≈ 13,8)'),
    ('Thông tin xuất bản', '2024, Vol. 7(7), e2419881, 12 trang'),
    ('DOI', '10.1001/jamanetworkopen.2024.19881'),
    ('Loại NC', 'Phân tích thứ cấp tiền đăng ký (prespecified secondary analysis) của RCT cụm (cluster RCT)'),
    ('Mẫu', '89 gia đình (181 trẻ 4–17 tuổi) từ 10 thành phố miền Nam Đan Mạch'),
    ('Đăng ký', 'ClinicalTrials.gov: NCT04098913'),
    ('Tài trợ', 'European Research Council (grant 716657); UK Medical Research Council'),
])
add_page_ref(doc, '1–12', 'JAMA Network Open', 'Vol. 7(7), 2024')

# TÓM TẮT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Tầm quan trọng: Sử dụng phương tiện màn hình quá mức đã được liên kết với SKTT kém hơn ở trẻ em và VTN trong nhiều NC quan sát. Tuy nhiên, bằng chứng thực nghiệm hỗ trợ giả thuyết này còn THIẾU.')
add_p(doc, 'Mục tiêu: Đánh giá tác động của can thiệp giảm sử dụng phương tiện màn hình giải trí trong 2 tuần lên SKTT trẻ em và VTN.')

p = doc.add_paragraph()
r = p.add_run('Thiết kế, bối cảnh và người tham gia: Phân tích thứ cấp tiền đăng ký của RCT cụm với theo dõi 2 tuần, bao gồm 89 gia đình (181 trẻ em và VTN) từ 10 thành phố Đan Mạch, miền Nam. Tất cả quy trình thực hiện tại nhà người tham gia. Tuyển chọn từ 06/06/2019 đến 30/03/2021. Phân tích thực hiện 01/01–30/11/2023.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
r = p.add_run('Can thiệp: Gia đình được phân ngẫu nhiên vào nhóm giảm screen time hoặc nhóm đối chứng. Can thiệp giảm screen time 2 tuần: giảm thời gian giải trí xuống ≤3 giờ/tuần/người, nộp smartphone và tablet cho nhà nghiên cứu.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
r = p.add_run('Kết quả chính: Chênh lệch trung bình giữa 2 nhóm trong thay đổi tổng khó khăn SDQ: −1,67 (KTC 95%: −2,68 đến −0,67), Cohen d = 0,53 (trung bình). Tác động mạnh nhất trên triệu chứng NỘI HÓA (cảm xúc + bạn bè): −1,03 (KTC 95%: −1,76 đến −0,29). Hành vi xã hội tích cực cũng cải thiện: 0,84 (KTC 95%: 0,39–1,30). Triệu chứng ngoại hóa KHÔNG cải thiện đáng kể: −0,55 (KTC 95%: −1,20 đến 0,10). Tuân thủ can thiệp: 97% (83/86 trẻ).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_p(doc, 'Kết luận: Khi toàn bộ gia đình — cả cha mẹ, trẻ em và VTN — giảm sử dụng phương tiện màn hình giải trí trong 2 tuần, điều này có thể tác động tích cực lên SKTT trẻ em/VTN. Can thiệp nhắm vào giảm tổng thể screen time giải trí, không chỉ ra loại hoạt động cụ thể.')

# ĐÁNH GIÁ NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'JAMA Network Open Q1 IF ≈ 13,8 — tạp chí y khoa hàng đầu thế giới.',
    'RCT — bằng chứng NHÂN QUẢ (cao nhất trong kim tự tháp bằng chứng). Đây là RCT ĐẦU TIÊN dựa trên gia đình về giảm screen time và SKTT trẻ em.',
    'Cohen d = 0,53 (trung bình) — kích thước hiệu ứng CÓ Ý NGHĨA LÂM SÀNG. Tương đương với can thiệp CBT đa thành phần (Bennett et al. 2024, Lancet; Chorpita et al. 2017, JCCP).',
    'Tác động mạnh nhất trên NỘI HÓA (cảm xúc + bạn bè) — phù hợp lo âu. Nature 2025 (QT27, Fassi et al.): VTN nội hóa nhạy cảm MXH hơn (g = 0,46).',
    'Tuân thủ 97% — can thiệp khả thi. Đo KHÁCH QUAN bằng tracker smartphone + tablet + PC + TV monitor.',
    'Tiền đăng ký ClinicalTrials.gov — giảm thiên lệch báo cáo chọn lọc.',
    'Can thiệp TOÀN GIA ĐÌNH — cả cha mẹ cùng giảm. Mô hình phù hợp cho VN (gia đình quan trọng).',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Mẫu NHỎ: 89 gia đình (181 trẻ). Sức mạnh thống kê hạn chế cho phân tích tiểu nhóm.',
    'Can thiệp CHỈ 2 tuần — hiệu quả dài hạn chưa rõ. Li 2025 (QT22): tác động dọc 12 tháng YẾU.',
    'Chỉ Đan Mạch — quốc gia giàu, Bắc Âu, văn hóa khác VN rất lớn.',
    'Mở nhãn (open-label) — gia đình BIẾT mình ở nhóm nào → thiên lệch kỳ vọng. Cha mẹ báo cáo SDQ có thể bị ảnh hưởng.',
    'Phân tích thứ cấp — KHÔNG phải kết quả chính của RCT gốc (kết quả chính: hoạt động thể chất).',
    'SDQ đo bởi PHỤ HUYNH — phản ánh nhận thức cha mẹ, không phải trẻ tự báo cáo.',
    'Gia đình TÌNH NGUYỆN — có thể có động cơ cao, không đại diện dân số chung.',
    'Nhóm đối chứng cũng giảm screen time một phần — làm giảm tương phản, có thể ĐÁNH GIÁ THẤP tác động.',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hướng cải thiện:', bold=True)
for b in [
    'RCT lớn hơn (200+ gia đình) với theo dõi ≥3 tháng.',
    'Phân biệt LOẠI screen time (MXH, game, học tập, TV) — Li 2025 (QT22) và Nature 2025 (QT27) cũng khuyến nghị.',
    'RCT tại trường VN — kết hợp giảm screen time + CBT nhóm + PE (BMC NMA 2025, QT29: cả 3 hiệu quả).',
    'Đo bằng trẻ TỰ BÁO CÁO (PHQ-A, CAS-8) thay vì chỉ phụ huynh (SDQ).',
]:
    add_p(doc, f'• {b}')

# 1. GIỚI THIỆU
add_page_ref(doc, '1–2', 'JAMA Network Open', '2024')
add_heading(doc, '1. GIỚI THIỆU', 2)

add_p(doc, 'Tỷ lệ trẻ em và VTN có SKTT kém đã tăng trong những năm gần đây ở nhiều quốc gia. Tại Mỹ, 29% VTN báo cáo SKTT kém trong tháng qua (Youth Risk Behavior Survey 2021). Dựa trên dữ liệu từ 45 nước châu Âu, 25% VTN 11–15 tuổi báo cáo triệu chứng liên quan sức khỏe tâm lý (lo lắng, kích ứng, khó ngủ). Cá nhân đặc biệt dễ bị tổn thương bởi trải nghiệm tiêu cực/căng thẳng trong giai đoạn VTN. Hơn nữa, NC cho thấy hài lòng cuộc sống tự báo cáo giảm mạnh nhất trong giai đoạn VTN so với các giai đoạn khác.')

add_p(doc, 'Những thập kỷ qua mang lại tiến bộ lớn về công nghệ số, và ngày càng nhiều trẻ em/VTN sở hữu thiết bị màn hình cá nhân. Sử dụng phương tiện màn hình đã trở thành khía cạnh trung tâm trong cuộc sống hàng ngày, với vô số lựa chọn giải trí: xem video, chơi game, tương tác MXH, giao tiếp gia đình và bạn bè. Sử dụng screen media là khái niệm phức tạp và rộng, bao gồm tương tác chủ động và thụ động với nhiều loại thiết bị và nội dung.')

add_p(doc, 'Mặc dù sử dụng thiết bị màn hình, đặc biệt smartphone, có thể hỗ trợ nhiều hoạt động hàng ngày và tương tác xã hội, lo ngại đã tăng về tác động tiêu cực tiềm năng lên SKTT trẻ em/VTN, nhưng NC vẫn thưa thớt và chưa kết luận. Một số NC quan sát tìm thấy mối liên quan kích thước trung bình giữa mức sử dụng màn hình cao và SKTT kém (Twenge & Campbell, 2019; Stiglic & Viner, 2019), trong khi một số nhà nghiên cứu cho rằng mối liên quan quan sát được là không đáng kể (Orben & Przybylski, 2019). Phân tích tổng hợp gần đây của Eirich et al. (2022) tìm thấy mối liên quan đáng kể giữa screen time và vấn đề nội hóa/ngoại hóa ở trẻ ≤12 tuổi.')

add_p(doc, 'Tuy nhiên, NC quan sát KHÔNG THỂ xác lập nhân quả — có thể SKTT kém dẫn đến tăng screen time, hoặc có yếu tố giao thoa. RCT là thiết kế duy nhất có thể kiểm tra nhân quả. Trước nghiên cứu này, có 2 RCT nhỏ nhắm vào giảm MXH ở sinh viên/người trưởng thành — chưa có RCT nào ở trẻ em/VTN và dựa trên gia đình.')

# 2. PHƯƠNG PHÁP
add_page_ref(doc, '2–5', 'JAMA Network Open', '2024')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)

add_p(doc, '2.1. Thiết kế thử nghiệm', bold=True)
add_p(doc, 'RCT cụm (cluster RCT) — đơn vị ngẫu nhiên hóa là GIA ĐÌNH (không phải cá nhân). Thuộc thử nghiệm SCREENS (Screen Media Reduction Experiment for Families). Tiền đăng ký: ClinicalTrials.gov NCT04098913. Phê duyệt đạo đức: Ủy ban Đạo đức Miền Nam Đan Mạch. Đồng ý bằng văn bản từ cha mẹ; trẻ em có mặt và có thể hỏi câu hỏi; dấu hiệu không đồng ý của trẻ được coi là chống chỉ định tham gia.')

add_p(doc, '2.2. Người tham gia', bold=True)
add_p(doc, 'Tuyển từ khảo sát dân số về hành vi screen media gia đình (Rasmussen et al. 2020). Thư mời khảo sát gửi đến cha mẹ có ≥1 trẻ 6–10 tuổi tại 10 thành phố miền Nam Đan Mạch. Người tham gia được chọn ngẫu nhiên bởi Cơ quan Dữ liệu Y tế Đan Mạch từ hệ thống đăng ký dân sự.')
add_p(doc, 'Tiêu chuẩn chọn: (1) Cha mẹ có screen time giải trí > phân vị 40 (dựa trên 1.000 phản hồi đầu tiên); (2) Cha mẹ làm việc hoặc học toàn thời gian; (3) Không có trẻ < 4 tuổi trong hộ; (4) Ít nhất 1 trẻ + 1 người lớn sẵn sàng tham gia và nộp smartphone/tablet.')
add_p(doc, 'Quy trình: 1.420 gia đình quan tâm → 408 đủ điều kiện → 92 đồng ý + hoàn thành baseline → 3 bỏ cuộc → 89 gia đình (181 trẻ) được phân ngẫu nhiên: 45 gia đình can thiệp (86 trẻ), 44 gia đình đối chứng (95 trẻ).')

add_p(doc, '2.3. Can thiệp', bold=True)
add_p(doc, 'Nhóm can thiệp: Giảm screen time giải trí xuống ≤3 giờ/tuần/người trong 2 tuần. NỘP smartphone và tablet cho nhà nghiên cứu. Cho phép xem TV tối đa 1 tập/ngày cùng gia đình. Cho phép dùng máy tính cho công việc/học tập (có phần mềm theo dõi).')
add_p(doc, 'Nhóm đối chứng: Tiếp tục thói quen screen media bình thường cho đến khi hoàn thành đo lường theo dõi, sau đó được mời thử can thiệp.')

add_p(doc, '2.4. Tuân thủ can thiệp', bold=True)
add_p(doc, 'Đo KHÁCH QUAN bằng: (1) Ứng dụng tracker trên smartphone và tablet do nhà nghiên cứu cài đặt; (2) Phần mềm PC; (3) Thiết bị giám sát TV. Kết quả: 97% (83/86) trẻ can thiệp tuân thủ đủ (≤7 giờ/tuần — cho phép chênh lệch nhỏ so với mục tiêu 3 giờ).')

add_p(doc, '2.5. Biến kết quả', bold=True)
add_p(doc, 'Bảng hỏi Điểm mạnh và Khó khăn (SDQ — Strengths and Difficulties Questionnaire): 25 mục, phụ huynh trả lời. 5 tiểu thang (5 mục mỗi tiểu thang): (1) Triệu chứng cảm xúc, (2) Vấn đề bạn bè, (3) Vấn đề hành vi, (4) Tăng động, (5) Hành vi xã hội tích cực. Tổng khó khăn = (1)+(2)+(3)+(4). Nội hóa = (1)+(2). Ngoại hóa = (3)+(4). Baseline: cha mẹ báo cáo 6 tháng trước. Follow-up: 2 tuần sau (cuối can thiệp). Phiên bản Đan Mạch phù hợp tuổi (2–4, 5–6, 4–10, 11–17).')

add_p(doc, '2.6. Phân tích', bold=True)
add_p(doc, 'Mô hình hồi quy tobit hỗn hợp (mixed-effects tobit regression) — xử lý biến kết quả có nhiều giá trị 0 (floor effect). Hiệu ứng ngẫu nhiên: gia đình (clustering). Điều chỉnh cho tuổi. Intention-to-treat (ITT). Sensitivity analysis: multiple imputation bằng chained equations cho dữ liệu thiếu. Phân tích tiểu nhóm hậu hoc: giới, tuổi (4–7, 8–10, 11–17), mức screen time baseline, mức tổng khó khăn baseline. Stata 18. P < 0,05 hai phía.')

# 3. KẾT QUẢ
add_page_ref(doc, '5–8', 'JAMA Network Open', '2024')
add_heading(doc, '3. KẾT QUẢ', 2)

add_p(doc, '3.1. Đặc điểm mẫu', bold=True)

add_heading(doc, 'Bảng 1. Đặc điểm người tham gia tại baseline', 3)
add_table(doc,
    ['Đặc điểm', 'Nhóm can thiệp (45 GĐ, 86 trẻ)', 'Nhóm đối chứng (44 GĐ, 95 trẻ)'],
    [['Tuổi TB (SD)', '8,6 (2,7)', '9,5 (2,5)'],
     ['Nữ', '42 (49%)', '57 (60%)'],
     ['Nam', '44 (51%)', '38 (40%)'],
     ['Tổng khó khăn SDQ — trung vị (IQR)', '7 (4–10)', '7 (4–10)'],
     ['Nội hóa — trung vị (IQR)', '3 (1–5)', '2 (1–5)'],
     ['Ngoại hóa — trung vị (IQR)', '4 (2–5)', '3 (2–5)'],
     ['Dữ liệu thiếu SDQ baseline', '4%', '—'],
     ['Dữ liệu thiếu SDQ follow-up', '11% (9 can thiệp, 10 đối chứng)', '—']],
    widths=[5.0, 4.5, 4.5])
add_p(doc, 'Đặc điểm baseline tương tự giữa 2 nhóm. Hầu hết trẻ nằm trong phạm vi bình thường cho SDQ trẻ Đan Mạch 5–17 tuổi.', size=9, italic=True)

add_p(doc, '3.2. Tác động can thiệp lên SKTT', bold=True)

add_heading(doc, 'Bảng 2. Tác động can thiệp giảm screen time lên SKTT (SDQ)', 3)
add_table(doc,
    ['Chỉ số SDQ', 'Chênh lệch (can thiệp − đối chứng)', 'KTC 95%', 'Cohen d', 'p', 'Ý nghĩa'],
    [['Tổng khó khăn', '−1,67', '−2,68 đến −0,67', '0,53', '<0,05', 'CẢI THIỆN — TRUNG BÌNH'],
     ['Nội hóa (cảm xúc + bạn bè)', '−1,03', '−1,76 đến −0,29', '—', '<0,05', 'CẢI THIỆN — MẠNH NHẤT'],
     ['Ngoại hóa (hành vi + tăng động)', '−0,55', '−1,20 đến 0,10', '—', '>0,05', 'KHÔNG đáng kể'],
     ['Hành vi xã hội tích cực', '0,84', '0,39 đến 1,30', '—', '<0,05', 'CẢI THIỆN'],
     ['Sensitivity (multiple imputation)', '−1,71', '−3,10 đến −0,33', '—', '<0,05', 'Kết quả tương tự']],
    widths=[4.0, 3.5, 3.0, 1.5, 1.0, 3.0])
add_p(doc, 'Chênh lệch 1,67 điểm SDQ tương đương hiệu quả can thiệp CBT đa thành phần (Bennett et al. 2024, Lancet: chênh 1,7 điểm; Chorpita et al. 2017: chênh 1,5 điểm). Cohen d = 0,53 > ngưỡng trung bình (0,50).', size=9, italic=True)

add_p(doc, '3.3. Phân tích tiểu nhóm (post hoc)', bold=True)

add_heading(doc, 'Bảng 3. Phân tích tiểu nhóm — Tổng khó khăn SDQ', 3)
add_table(doc,
    ['Tiểu nhóm', 'Chênh lệch (ước tính)', 'Tương tác p', 'Ghi chú'],
    [['Nam', 'Lớn hơn nữ', '>0,05 (không đáng kể)', 'Xu hướng nam đáp ứng tốt hơn'],
     ['Nữ', 'Nhỏ hơn nam', '', ''],
     ['4–7 tuổi', 'Trung bình', '>0,05', ''],
     ['8–10 tuổi', 'Lớn hơn', '>0,05', 'Xu hướng 8–10 đáp ứng tốt nhất'],
     ['11–17 tuổi', 'Nhỏ hơn', '>0,05', ''],
     ['Screen time baseline CAO', 'Lớn hơn', '>0,05', 'Xu hướng: giảm nhiều → cải thiện nhiều'],
     ['SDQ baseline > trung vị', 'LỚN hơn', '>0,05', 'CÓ ý nghĩa lâm sàng — nhóm nguy cơ cao'],
     ['SDQ baseline ≤ trung vị', 'KHÔNG đáng kể', '', 'Sàn (floor effect) — ít chỗ cải thiện']],
    widths=[4.0, 3.0, 3.0, 4.0])
add_p(doc, 'Không có tương tác tiểu nhóm đạt ý nghĩa thống kê — mẫu nhỏ hạn chế sức mạnh. Tuy nhiên, xu hướng gợi ý can thiệp hiệu quả hơn ở: nam, 8–10 tuổi, screen time cao, SDQ khó khăn cao.', size=9, italic=True)

# 4. THẢO LUẬN
add_page_ref(doc, '7–9', 'JAMA Network Open', '2024')
add_heading(doc, '4. THẢO LUẬN', 2)

add_p(doc, 'Kết quả RCT cụm cho thấy giảm screen time giải trí trong gia đình 2 tuần có tác động tích cực đáng kể lên SKTT trẻ em/VTN. Xem xét các tiểu lĩnh vực tâm thần, giảm screen time cải thiện triệu chứng nội hóa (cảm xúc + bạn bè) và hành vi xã hội tích cực. Phát hiện gợi ý: lợi ích rõ rệt nhất của giảm screen time là giảm vấn đề nội hóa và tăng tương tác xã hội tích cực.')

add_p(doc, 'Theo hiểu biết của chúng tôi, đây là RCT dựa trên gia đình ĐẦU TIÊN kiểm tra tác động giảm screen time giải trí lên SKTT trẻ em/VTN. NC xác nhận giảm screen time cải thiện nhiều khía cạnh SKTT trong ngắn hạn, phù hợp với NC quan sát (Eirich et al. 2022, JAMA Psychiatry; Stiglic & Viner, 2019). Trái với NC trước báo cáo kích thước hiệu ứng không đáng kể (Orben & Przybylski, 2019), kết quả RCT cho thấy hiệu ứng TRUNG BÌNH (d = 0,53). Chênh lệch 1,67 điểm SDQ tương đương hiệu quả can thiệp tâm lý đa thành phần cho SKTT trẻ em (Bennett et al. 2024, Lancet).')

add_p(doc, 'Tác động mạnh nhất trên nội hóa:', bold=True)
add_p(doc, 'Phù hợp với Nature 2025 (QT27, Fassi et al.): VTN có rối loạn nội hóa (lo âu, trầm cảm) nhạy cảm hơn với MXH (g = 0,46 vs ngoại hóa g = 0,38). Cơ chế tiềm năng: screen time thay thế tương tác xã hội trực tiếp → giảm cơ hội phát triển kỹ năng xã hội → tăng vấn đề bạn bè và cảm xúc. Khi giảm screen time, trẻ em TĂNG tương tác trực tiếp → cải thiện.')

add_p(doc, 'Tiểu nhóm:', bold=True)
add_p(doc, 'Xu hướng (không đạt ý nghĩa thống kê): nam, 8–10 tuổi, screen time cao baseline có thể đáp ứng tốt hơn. Khác biệt có thể do loại tương tác screen media khác nhau theo giới và tuổi. Đáng chú ý: can thiệp hiệu quả hơn ở trẻ có SDQ khó khăn cao baseline — nhóm NGUY CƠ CAO hưởng lợi nhiều nhất. Trong nhóm SDQ thấp (bình thường), CAN THIỆP KHÔNG CÓ TÁC ĐỘNG — có thể do floor effect (đã ở mức tối ưu).')

add_p(doc, 'Can thiệp toàn gia đình:', bold=True)
add_p(doc, 'Điểm mạnh độc đáo — CẢ gia đình giảm screen time cùng nhau. Có thể: (1) tạo môi trường hỗ trợ (trẻ không bị cô lập khi gia đình vẫn dùng), (2) tăng thời gian chất lượng gia đình, (3) cha mẹ làm gương. Tuy nhiên, NC tương lai cần xác nhận liệu sự tham gia gia đình là thành phần THEN CHỐT hay không.')

add_p(doc, 'Cơ chế tiềm năng:', bold=True)
add_p(doc, 'Chỉ có thể suy đoán: (1) Giảm screen time → tăng tương tác trực tiếp với bạn bè, gia đình; (2) Cải thiện giấc ngủ (Zhu 2025, QT05: giấc ngủ <5h AOR = 13,71 cho trầm cảm); (3) Tăng hoạt động thể chất (BMC NMA 2025, QT29: PE hiệu quả cho lo âu); (4) Giảm so sánh xã hội trực tuyến (Nature 2025, QT27: g = 0,30 cho so sánh xã hội).')

# 5. HẠN CHẾ
add_heading(doc, '5. ĐIỂM MẠNH VÀ HẠN CHẾ', 2)

add_p(doc, 'Điểm mạnh:', bold=True)
add_p(doc, '• RCT — thiết kế cao nhất. Đo tuân thủ KHÁCH QUAN (tracker, phần mềm, TV monitor) — khác biệt với tự báo cáo.')
add_p(doc, '• Tỷ lệ bỏ cuộc thấp và dữ liệu thiếu ít — giảm thiên lệch mất mẫu.')
add_p(doc, '• Tiền đăng ký — giảm thiên lệch báo cáo chọn lọc.')
add_p(doc, '• Can thiệp toàn gia đình — độc đáo và thực tế.')

add_p(doc, 'Hạn chế:', bold=True)
add_p(doc, '• Mở nhãn (open-label) — cha mẹ biết nhóm → có thể thiên lệch khi trả lời SDQ. Tương lai cần thiết kế mù (khó khả thi cho can thiệp hành vi).')
add_p(doc, '• Nhóm đối chứng cũng giảm screen time một phần → giảm tương phản → có thể ĐÁNH GIÁ THẤP tác động thực.')
add_p(doc, '• Gia đình tình nguyện — động cơ cao, không đại diện dân số chung.')
add_p(doc, '• SDQ phiên bản phụ huynh — phản ánh nhận thức cha mẹ nhiều hơn trải nghiệm trẻ.')
add_p(doc, '• Chỉ 2 tuần — hiệu quả dài hạn chưa rõ. Có thể trẻ quay lại mức screen time cũ sau can thiệp.')
add_p(doc, '• Mẫu nhỏ (181 trẻ) — hạn chế phân tích tiểu nhóm, cần RCT lớn hơn đặc biệt ở nhóm nguy cơ cao.')

# 6. KẾT LUẬN
add_heading(doc, '6. KẾT LUẬN', 2)
add_p(doc, 'Tổng hợp, kết quả phân tích thứ cấp RCT cho thấy khi toàn bộ gia đình — cả cha mẹ, trẻ em và VTN — giảm sử dụng phương tiện màn hình giải trí trong 2 tuần, điều này có thể tác động tích cực lên điểm mạnh và khó khăn hành vi của trẻ em/VTN (Cohen d = 0,53). Tác động mạnh nhất trên triệu chứng nội hóa (cảm xúc + bạn bè) — phù hợp với lo âu. Can thiệp nhắm giảm tổng thể screen time giải trí, không chỉ ra loại hoạt động cụ thể.')
add_p(doc, 'NC tương lai cần: (1) khám phá tác động khác biệt của các loại screen time (game, MXH, video); (2) xác nhận liệu sự tham gia gia đình là thành phần then chốt; (3) thử nghiệm ở nhóm nguy cơ cao; (4) theo dõi dài hạn ≥3 tháng; (5) RCT tại các nước LMIC bao gồm VN.')

# TÀI LIỆU THAM KHẢO
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
refs = [
    'Schmidt-Persson, J., et al. (2024). Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial. JAMA Network Open, 7(7), e2419881.',
    'Eirich, R., et al. (2022). Association of screen time with internalizing and externalizing behavior problems in children 12 years or younger. JAMA Psychiatry, 79(5), 393–405.',
    'Orben, A. & Przybylski, A.K. (2019). The association between adolescent well-being and digital technology use. Nature Human Behaviour, 3(2), 173–182.',
    'Stiglic, N. & Viner, R.M. (2019). Effects of screentime on the health and well-being of children and adolescents. BMJ Open, 9(1), e023191.',
    'Twenge, J.M. & Campbell, W.K. (2019). Media use is linked to lower psychological well-being. Psychiatric Quarterly, 90(2), 311–331.',
    'Rasmussen, M.G.B., et al. (2020). Short-term efficacy of reducing screen media use on physical activity, sleep, and physiological stress: The SCREENS RCT protocol. BMC Public Health, 20(1), 380.',
    'Bennett, S.D., et al. (2024). Clinical effectiveness of psychological therapy for children with epilepsy. Lancet, 403, 1254–1266.',
    'Pedersen, J., et al. (2022). Effects of reducing screen media use on physical activity among families with children. IJBNPA, 19, 1–12.',
    'Fassi, L., et al. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9, 1283–1299.',
    'Li, S.H., et al. (2025). Cross-sectional and longitudinal associations of screen time with adolescent depression and anxiety. BJCP, 64, 873–887.',
]
for ref in refs:
    add_p(doc, ref, size=10)
add_p(doc, '(Xem danh mục đầy đủ trong bài gốc — 44 tài liệu tham khảo)', size=10, italic=True)

# BẢNG VIẾT TẮT
add_abbreviation_table(doc, [
    ('SDQ', 'Strengths and Difficulties Questionnaire — Bảng hỏi Điểm mạnh và Khó khăn'),
    ('RCT', 'Randomized Clinical Trial — Thử nghiệm Lâm sàng Ngẫu nhiên'),
    ('ITT', 'Intention-to-Treat — Phân tích Ý định Điều trị'),
    ('IQR', 'Interquartile Range — Khoảng Tứ phân vị'),
    ('KTC', 'Khoảng Tin cậy (Confidence Interval)'),
    ('ST', 'Screen Time — Thời gian Sử dụng Màn hình'),
    ('MXH', 'Mạng xã hội'),
    ('SKTT', 'Sức khỏe tâm thần'),
    ('VTN', 'Vị thành niên'),
    ('CBT', 'Cognitive-Behavioral Therapy — Liệu pháp Nhận thức-Hành vi'),
    ('PE', 'Physical Exercise — Hoạt động Thể chất'),
    ('CONSORT', 'Consolidated Standards of Reporting Trials'),
    ('SCREENS', 'Screen Media Reduction Experiment for Families'),
    ('NMA', 'Network Meta-Analysis — Phân tích Tổng hợp Mạng'),
])

# PHẢN BIỆN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')

add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'JAMA Network Open Q1 IF ≈ 13,8 — tạp chí y khoa hàng đầu. Tiền đăng ký ClinicalTrials.gov.',
    'RCT — thiết kế CAO NHẤT cho bằng chứng nhân quả. ĐẦU TIÊN dựa trên gia đình cho screen time → SKTT trẻ em.',
    'Cohen d = 0,53 (trung bình) — CÓ Ý NGHĨA LÂM SÀNG. Tương đương CBT đa thành phần (Bennett et al. 2024, Lancet: 1,7 điểm; Chorpita et al. 2017: 1,5 điểm). AJP 2024 (QT28, Zugman et al.): CBT đơn đáp ứng 59,7%, tương đương mức cải thiện này.',
    'Tác động mạnh nhất trên NỘI HÓA — phù hợp lo âu. Xác nhận Nature 2025 (QT27, Fassi et al.): VTN nội hóa nhạy cảm MXH hơn (g = 0,46). Norway 2025 (QT21): MXH giải thích MỘT PHẦN xu hướng tăng.',
    'Đo tuân thủ KHÁCH QUAN (tracker app, PC software, TV monitor) — vượt trội so với tự báo cáo. Li 2025 (QT22) và Nature 2025 (QT27) đều dùng tự báo cáo.',
    'Can thiệp TOÀN GIA ĐÌNH — độc đáo. Phù hợp văn hóa VN (gia đình quan trọng). 59 Countries 2025 (QT31): cha mẹ kiểm tra bài AOR = 0,75 (bảo vệ).',
    'Tuân thủ 97% — can thiệp khả thi mặc dù đòi hỏi nộp smartphone.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'Mẫu RẤT NHỎ: 89 gia đình (181 trẻ). Hạn chế sức mạnh thống kê cho tiểu nhóm. JAACAP 2024 (QT23): 13,7 TRIỆU hồ sơ. Norway 2025 (QT21): 979K. Cần RCT lớn hơn nhiều.',
    'Can thiệp CHỈ 2 tuần — hiệu quả dài hạn CHƯA RÕ. Li 2025 (QT22): tác động dọc 12 tháng YẾU (b = 0,15, p = 0,007 cho trầm cảm; KHÔNG đáng kể cho lo âu). Có thể hiệu ứng biến mất khi trẻ quay lại screen time.',
    'Chỉ Đan Mạch — giàu, Bắc Âu, dân số nhỏ. Văn hóa screen media RẤT KHÁC VN (Zalo, TikTok vs Instagram). Hoàng Trung Học 2025 (VN14): game điện tử Beta = 0,176 ở VTN VN.',
    'Mở nhãn — cha mẹ BIẾT nhóm → thiên lệch kỳ vọng khi trả lời SDQ. Cần thiết kế mù hoặc đo trẻ tự báo cáo (PHQ-A, CAS-8 như Li 2025, QT22).',
    'SDQ phiên bản PHỤ HUYNH — phản ánh nhận thức cha mẹ, KHÔNG phải trải nghiệm trẻ. V-NAMHS 2022 (VN02): chỉ 5,1% phụ huynh VN nhận ra con cần giúp đỡ — cha mẹ có thể không nhạy cảm.',
    'Phân tích THỨ CẤP — kết quả chính RCT gốc là hoạt động thể chất (Pedersen et al. 2022), KHÔNG phải SKTT.',
    'Nhóm đối chứng CŨNG giảm screen time một phần → tương phản giảm → tác động có thể bị ĐÁNH GIÁ THẤP.',
    'Gia đình TÌNH NGUYỆN — có động cơ cao hơn dân số chung. Không đại diện gia đình VN (nơi áp lực học tập và ít thời gian giải trí).',
    'KHÔNG phân biệt LOẠI screen time (MXH, game, TV, học tập) — can thiệp giảm TẤT CẢ. Nature 2025 (QT27): cần phân biệt vì tác động khác nhau. Hoàng Trung Học 2025 (VN14): game β = 0,176, MXH có thể khác.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống nghiên cứu / Research Gap:', bold=True)
for s in [
    'VN CHƯA CÓ BẤT KỲ RCT nào về giảm screen time ở trẻ em/VTN — GAP CỰC LỚN. Đây là ưu tiên #7 trong cross-study.',
    'Cần RCT giảm screen time TẠI TRƯỜNG VN — mô hình thực tế hơn gia đình (cha mẹ VN bận, khó tham gia). Kết hợp với giảm áp lực học tập (Vĩnh Lộc 2024, VN20: ESSA ≥59 tăng nguy cơ).',
    'So sánh can thiệp PHÂN BIỆT LOẠI screen time: giảm MXH vs giảm game vs giảm TV — tác động khác nhau? Hoàng Trung Học 2025 (VN14) gợi ý game riêng biệt.',
    'Kết hợp giảm screen time + CBT nhóm + PE tại trường — BMC NMA 2025 (QT29): cả CBT (SUCRA 0,66) VÀ PE (0,51) hiệu quả cho lo âu. AJP 2024 (QT28): CBT+SSRI 80,7% đáp ứng.',
    'RCT dài hạn (≥3 tháng) — Li 2025 (QT22) cho thấy tác động 12 tháng YẾU. Cần xác nhận hiệu quả bền vững.',
    'Đo trẻ TỰ BÁO CÁO (PHQ-A + CAS-8 + SIAS-17) thay vì chỉ SDQ phụ huynh. Jefferies 2020 (QT35): 18% "false negatives" — trẻ có lo âu nhưng không nhận ra (phụ huynh càng không).',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '33_JAMA_ScreenMedia_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
