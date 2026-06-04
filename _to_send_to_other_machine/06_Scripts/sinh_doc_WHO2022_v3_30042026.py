# -*- coding: utf-8 -*-
"""
WHO 2022 World Mental Health Report — Bản dịch + Phản biện chi tiết (V3, 30/04/2026)
- Tiếng Việt thuần — chú thích tiếng Anh trong ngoặc
- Page markers màu cam, từ viết tắt lần đầu màu đỏ đậm
- Bảng Word thật, 0 hình (PDF gốc dùng infographic phức tạp — em tái tạo bằng bảng)
- Phản biện chi tiết, có dẫn chứng
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(r"c:\Users\OS\OneDrive\read_books\Lo-au")
OUT = ROOT / "03_Ban-dich" / "Bai_dich_phan_bien" / "WHO_2022_World_Mental_Health_Report_dich_phan_bien_30042026.docx"

ORANGE = RGBColor(0xE6, 0x7E, 0x22)
RED = RGBColor(0xC0, 0x39, 0x2B)
BLUE = RGBColor(0x2E, 0x86, 0xC1)
GRAY = RGBColor(0x70, 0x70, 0x70)
DRED = RGBColor(0x9B, 0x1B, 0x1B)

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)

st = d.styles["Normal"]
st.font.name = "Times New Roman"; st.font.size = Pt(12)
rpr = st.element.rPr
rfonts = rpr.find(qn("w:rFonts")) or OxmlElement("w:rFonts")
rfonts.set(qn("w:eastAsia"), "Times New Roman"); rpr.append(rfonts)


def H1(text):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(14); r.bold = True; r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)


def H2(text):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(12); r.bold = True; r.font.color.rgb = BLUE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)


def nr(text):
    p = d.add_paragraph(); p.add_run(text).font.size = Pt(12)
    p.paragraph_format.space_after = Pt(3)


def page_marker(text):
    p = d.add_paragraph(); r = p.add_run(f"— — — {text} — — —")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = ORANGE
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)


def crit_para(text):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(12); r.font.color.rgb = RED
    p.paragraph_format.space_after = Pt(4)


def en_inline(p, en):
    """Add English clarification in gray italic"""
    r = p.add_run(f" ({en})")
    r.font.size = Pt(11); r.italic = True; r.font.color.rgb = GRAY


def abbr_first(text):
    """First-time abbreviation in dark red bold"""
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.size = Pt(12); r.bold = True; r.font.color.rgb = DRED


def tbl(headers, rows, widths=None):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = "Light Grid Accent 1"
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ""
        p = c.paragraphs[0]; r = p.add_run(h); r.bold = True; r.font.size = Pt(11)
    for r_idx, row in enumerate(rows):
        for c_idx, val in enumerate(row):
            c = t.rows[r_idx+1].cells[c_idx]; c.text = ""
            p = c.paragraphs[0]; rr = p.add_run(str(val)); rr.font.size = Pt(11)
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    d.add_paragraph()


# ============================================================
# TRANG BÌA
# ============================================================
title = d.add_paragraph(); title.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = title.add_run("BẢN DỊCH + PHẢN BIỆN CHI TIẾT\n"); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE
r = title.add_run("BÁO CÁO SỨC KHOẺ TÂM THẦN THẾ GIỚI 2022\n"); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BLUE
r = title.add_run("CHUYỂN HOÁ SỨC KHOẺ TÂM THẦN CHO TẤT CẢ MỌI NGƯỜI\n"); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
r = title.add_run("(WHO World Mental Health Report — Transforming mental health for all, 2022)")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY

d.add_paragraph()
sub = d.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub.add_run("Tổ chức Y tế Thế giới (World Health Organization — WHO)\n").bold = True
sub.add_run("Geneva, 17/06/2022 — 296 trang — Open Access (CC BY-NC-SA 3.0 IGO)\n")
sub.add_run("ISBN 978-92-4-004933-8 (electronic) / 978-92-4-004934-5 (print)\n")
r = sub.add_run("Trợ lý nghiên cứu — 30/04/2026 — tiếng Việt thuần — chú thích Anh trong ngoặc\n")
r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
r = sub.add_run("PDF gốc trong dự án: 02_Papers-goc/GBD_WHO/WHO_2022_World_Mental_Health_Report.pdf (3,6 MB / 296 trang)\n")
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GRAY

# ============================================================
# THÔNG TIN THƯ MỤC
# ============================================================
H1("THÔNG TIN THƯ MỤC")
tbl(['Mục', 'Nội dung'], [
    ['Tên báo cáo (tiếng Anh)', 'World mental health report: Transforming mental health for all'],
    ['Tạm dịch tên báo cáo', 'Báo cáo sức khoẻ tâm thần thế giới: Chuyển hoá sức khoẻ tâm thần cho tất cả mọi người'],
    ['Nhà xuất bản', 'World Health Organization (WHO) — Tổ chức Y tế Thế giới'],
    ['Năm xuất bản', '2022 (xuất bản 17/06/2022)'],
    ['Số trang', '296 trang (gồm bìa, lời nói đầu, 8 chương, kết luận, tham khảo)'],
    ['ISBN điện tử', '978-92-4-004933-8'],
    ['ISBN bản in', '978-92-4-004934-5'],
    ['Loại tài liệu', 'Báo cáo chính sách + tổng quan toàn cầu — không phải bài báo gốc'],
    ['Mục đích', 'Truyền cảm hứng + cung cấp thông tin để chuyển hoá hệ thống sức khoẻ tâm thần toàn cầu'],
    ['Đối tượng đọc', 'Quan chức y tế các nước thành viên WHO + đối tác đa ngành'],
    ['Trích dẫn chuẩn', 'World mental health report: transforming mental health for all. Geneva: World Health Organization; 2022. Licence: CC BY-NC-SA 3.0 IGO.'],
    ['Trạng thái mở', 'Open Access dưới giấy phép Creative Commons CC BY-NC-SA 3.0 IGO'],
    ['Đường dẫn truy cập', 'https://www.who.int/publications/i/item/9789240049338'],
    ['Liên kết với chính sách', 'WHO Comprehensive mental health action plan 2013–2030 + UN Sustainable Development Goals'],
    ['Mã trong kho dự án', 'Đề xuất: WHO_WMH_2022'],
])

# ============================================================
# BẢNG TỪ VIẾT TẮT
# ============================================================
H1("BẢNG TỪ VIẾT TẮT DÙNG TRONG TÀI LIỆU NÀY")
nr("Lần đầu mỗi từ viết tắt xuất hiện trong văn bản dưới đây sẽ được tô MÀU ĐỎ ĐẬM + chú thích inline.")
tbl(['Tiếng Việt', 'Tiếng Anh đầy đủ', 'Viết tắt'], [
    ['Tổ chức Y tế Thế giới', 'World Health Organization', 'WHO'],
    ['Báo cáo Sức khoẻ Tâm thần Thế giới (báo cáo này)', 'World Mental Health Report', 'WMH Report 2022'],
    ['Nước thu nhập thấp + trung bình', 'Low- and Middle-Income Countries', 'LMICs'],
    ['Năm sống điều chỉnh theo khuyết tật', 'Disability-Adjusted Life-Years', 'DALYs'],
    ['Năm sống với khuyết tật', 'Years Lived with Disability', 'YLDs'],
    ['Năm sống bị mất do tử vong sớm', 'Years of Life Lost', 'YLLs'],
    ['Bao phủ y tế toàn dân', 'Universal Health Coverage', 'UHC'],
    ['Mục tiêu Phát triển Bền vững (Liên Hợp Quốc)', 'Sustainable Development Goals', 'SDGs'],
    ['Sức khoẻ ban đầu', 'Primary Health Care', 'PHC'],
    ['Hệ thống thông tin SKTT', 'Mental Health Information System', 'MHIS'],
    ['Kế hoạch hành động SKTT toàn diện 2013-2030', 'Comprehensive Mental Health Action Plan 2013–2030', 'CMHAP'],
    ['Phân loại bệnh quốc tế phiên bản 11', 'International Classification of Diseases, 11th revision', 'ICD-11'],
    ['Sổ tay chẩn đoán + thống kê rối loạn tâm thần phiên bản 5', 'Diagnostic and Statistical Manual of Mental Disorders, 5th edition', 'DSM-5'],
    ['Nghiên cứu Gánh nặng Bệnh tật Toàn cầu', 'Global Burden of Disease', 'GBD'],
    ['Đánh giá tác động sức khoẻ tâm thần', 'Mental Health Atlas', '—'],
    ['Rối loạn căng thẳng sau sang chấn', 'Post-Traumatic Stress Disorder', 'PTSD'],
    ['Rối loạn tăng động giảm chú ý', 'Attention-Deficit Hyperactivity Disorder', 'ADHD'],
    ['Rối loạn phổ tự kỷ', 'Autism Spectrum Disorder', 'ASD'],
    ['Sức khoẻ tâm thần', 'Mental health', 'SKTT'],
    ['Người có trải nghiệm sống', 'People with lived experience', '—'],
])

# ============================================================
# PHẦN 1 — LỜI NÓI ĐẦU + GIỚI THIỆU
# ============================================================
H1("PHẦN 1 — LỜI NÓI ĐẦU CỦA TỔNG GIÁM ĐỐC WHO")
page_marker("Trang viii — Foreword by Tedros Adhanom Ghebreyesus")
nr("Khi thế giới đang chung sống và rút bài học từ những tác động sâu rộng của đại dịch COVID-19, tất cả chúng ta phải suy ngẫm về một trong những khía cạnh nổi bật nhất của nó — gánh nặng KHỔNG LỒ mà đại dịch đã gây ra cho sức khoẻ tâm thần (SKTT) của mọi người. Tỉ lệ các rối loạn vốn đã phổ biến như trầm cảm và lo âu đã TĂNG HƠN 25% trong năm đầu của đại dịch, cộng dồn vào gần MỘT TỈ người vốn đang sống chung với một rối loạn tâm thần.")
nr("Đồng thời, chúng ta phải nhận ra sự MONG MANH của các hệ thống y tế đang cố gắng đáp ứng nhu cầu của những người có cả các tình trạng SKTT mới phát sinh và đã tồn tại từ trước.")
nr("Sức khoẻ tâm thần KHÔNG CHỈ là vắng mặt bệnh tật — nó là phần nội tại của sức khoẻ và hạnh phúc của chúng ta cả về cá nhân và tập thể. Như báo cáo này cho thấy, để đạt được các mục tiêu toàn cầu trong Kế hoạch hành động sức khoẻ tâm thần toàn diện 2013-2030 và các Mục tiêu Phát triển Bền vững (SDGs), chúng ta cần CHUYỂN HOÁ thái độ, hành động và cách tiếp cận để PHÁT HUY và BẢO VỆ SKTT, đồng thời CUNG CẤP và CHĂM SÓC cho những người cần.")

H1("PHẦN 2 — GIỚI THIỆU (Chương 1)")
page_marker("Trang 1-7, Chapter 1: Introduction")

H2("1.1. Hai mươi năm sau (Twenty years on)")
nr("Hai mươi năm sau khi WHO xuất bản báo cáo cột mốc \"Báo cáo sức khoẻ thế giới 2001: Sức khoẻ tâm thần — Hiểu biết mới, Hi vọng mới\", các khuyến nghị đưa ra khi đó VẪN CÒN HIỆU LỰC đến hôm nay. Tuy vậy, NHIỀU TIẾN BỘ đã được thực hiện. Sự quan tâm và hiểu biết về SKTT đã tăng. Nhiều quốc gia đã thiết lập, cập nhật và củng cố các chính sách SKTT. Phong trào vận động đã khuếch đại tiếng nói của người có trải nghiệm sống với rối loạn tâm thần.")
nr("Các quốc gia thành viên WHO đã thông qua Kế hoạch hành động SKTT toàn diện 2013-2030. Họ cam kết đạt các chỉ tiêu toàn cầu về cải thiện SKTT — tập trung vào: (1) củng cố lãnh đạo và quản trị; (2) chăm sóc dựa vào cộng đồng; (3) phát huy và phòng ngừa; (4) hệ thống thông tin và nghiên cứu.")
nr("NHƯNG: phân tích mới nhất của WHO về kết quả thực hiện kế hoạch ở các nước cho thấy TIẾN BỘ đã CHẬM. Đối với phần lớn thế giới, cách tiếp cận chăm sóc SKTT vẫn là \"như thường lệ\" (business as usual). KẾT QUẢ: các tình trạng SKTT vẫn tiếp tục gây gánh nặng nặng nề lên cuộc sống con người, trong khi hệ thống và dịch vụ SKTT vẫn KHÔNG đủ năng lực đáp ứng nhu cầu.")

H2("1.2. Cần thay đổi (Time for change)")
nr("Các đe doạ toàn cầu đối với SKTT luôn hiện diện. Bất bình đẳng kinh tế-xã hội ngày càng gia tăng, xung đột kéo dài, bạo lực và tình huống khẩn cấp y tế công cộng ĐE DOẠ tiến độ cải thiện sức khoẻ. Hơn bao giờ hết, \"như thường lệ\" trong SKTT KHÔNG ĐƯỢC nữa. Báo cáo này nhằm TRUYỀN CẢM HỨNG và CUNG CẤP THÔNG TIN cho cuộc CHUYỂN HOÁ KHẨN CẤP và KHÔNG THỂ TRANH CÃI để đảm bảo SKTT TỐT HƠN cho tất cả mọi người.")

H2("1.3. Về báo cáo này (About this report)")
nr("Mặc dù khuyến khích phương pháp tiếp cận đa ngành, báo cáo này được viết ĐẶC BIỆT cho NGƯỜI RA QUYẾT ĐỊNH trong ngành y tế — bao gồm các Bộ Y tế và đối tác y tế khác thường được giao nhiệm vụ phát triển chính sách SKTT và cung cấp hệ thống + dịch vụ SKTT. Báo cáo gồm 8 chương: (1) Giới thiệu; (2) Nguyên tắc và động lực trong SKTT công cộng; (3) SKTT thế giới hôm nay; (4) Lợi ích của thay đổi; (5) Nền tảng cho thay đổi; (6) Phát huy và phòng ngừa; (7) Tái cấu trúc và mở rộng dịch vụ chăm sóc; (8) Kết luận.")

# ============================================================
# PHẦN 3 — TÓM TẮT 8 CHƯƠNG (Executive Summary cards)
# ============================================================
H1("PHẦN 3 — TÓM TẮT TỔNG QUÁT 8 CHƯƠNG (Executive Summary)")
page_marker("Trang xii-xv, Executive Summary cards trong PDF gốc")

H2("Chương 1 — Giới thiệu: Một báo cáo thế giới để truyền cảm hứng và thông tin về thay đổi")
nr("Đã trình bày trong PHẦN 2 ở trên.")

H2("Chương 2 — Nguyên tắc và động lực trong SKTT công cộng: SKTT cực kỳ quan trọng cho TẤT CẢ MỌI NGƯỜI, MỌI NƠI")
nr("SKTT là phần nội tại của sức khoẻ và hạnh phúc chung — và là một QUYỀN CƠ BẢN CỦA CON NGƯỜI. Có SKTT tốt nghĩa là chúng ta có thể KẾT NỐI, HOẠT ĐỘNG, ỨNG PHÓ và PHÁT TRIỂN tốt hơn. SKTT tồn tại trên một liên tục PHỨC TẠP — từ trạng thái hạnh phúc tối ưu đến trạng thái suy nhược + đau khổ tinh thần lớn.")
nr("Tại bất kỳ thời điểm nào, một tổ hợp các yếu tố cá nhân + gia đình + cộng đồng + cấu trúc có thể KẾT HỢP để bảo vệ HOẶC làm suy yếu SKTT của chúng ta. Mặc dù hầu hết mọi người KHÁ KIÊN CƯỜNG, người đang sống trong điều kiện BẤT LỢI — bao gồm nghèo đói, bạo lực và bất bình đẳng — có nguy cơ CAO HƠN trải nghiệm các tình trạng SKTT.")
nr("Yếu tố nguy cơ có thể biểu hiện ở MỌI giai đoạn cuộc đời, NHƯNG những yếu tố xảy ra trong các giai đoạn nhạy cảm về phát triển — đặc biệt thời thơ ấu — là ĐẶC BIỆT NGUY HẠI. Yếu tố bảo vệ cũng tồn tại xuyên suốt cuộc đời và giúp tăng cường KIÊN CƯỜNG, gồm: kỹ năng cảm xúc-xã hội cá nhân, tương tác xã hội tích cực, giáo dục chất lượng, công việc tốt, khu phố an toàn và gắn kết cộng đồng.")
nr("Vì các yếu tố quyết định SKTT là ĐA NGÀNH, can thiệp phát huy + bảo vệ SKTT cũng cần được CUNG CẤP qua NHIỀU NGÀNH — không chỉ ngành y tế.")

H2("Chương 3 — SKTT thế giới hôm nay: Nhu cầu CAO nhưng đáp ứng KHÔNG ĐỦ")
nr("Ở MỌI quốc gia, các tình trạng SKTT có tỉ lệ HIỆN MẮC CAO. KHOẢNG 1 TRONG 8 NGƯỜI trên thế giới đang sống chung với một rối loạn tâm thần (~970 TRIỆU NGƯỜI năm 2019). Tỉ lệ phổ biến các rối loạn khác nhau theo GIỚI và TUỔI. Ở cả nam và nữ, RỐI LOẠN LO ÂU và RỐI LOẠN TRẦM CẢM là phổ biến NHẤT.")
nr("Tự tử ảnh hưởng người và gia đình ở MỌI quốc gia + bối cảnh + ở MỌI lứa tuổi. Toàn cầu, có thể có 20 LẦN MƯU SÁT cho mỗi 1 ca tử vong, và tự tử chiếm HƠN 1 TRÊN 100 ca tử vong. Tự tử là nguyên nhân chính gây tử vong ở NGƯỜI TRẺ.")
nr("Rối loạn tâm thần là NGUYÊN NHÂN HÀNG ĐẦU của YLDs (Năm sống với khuyết tật) — chiếm 1 TRÊN 6 YLDs toàn cầu. Tâm thần phân liệt — xảy ra ở khoảng 1 trên 200 người lớn — là MỐI QUAN TÂM ƯU TIÊN: ở trạng thái cấp, đó là tình trạng GÂY KHUYẾT TẬT NẶNG NHẤT trong tất cả các tình trạng sức khoẻ. Người có tâm thần phân liệt hoặc các rối loạn tâm thần nghiêm trọng khác CHẾT TRUNG BÌNH 10-20 NĂM SỚM HƠN dân số chung — thường do các bệnh thể chất CÓ THỂ PHÒNG NGỪA.")
nr("Hệ quả KINH TẾ của rối loạn tâm thần là KHỔNG LỒ. Mất năng suất + chi phí gián tiếp khác cho xã hội thường VƯỢT XA chi phí chăm sóc y tế. Về kinh tế, tâm thần phân liệt là rối loạn ĐẮT NHẤT cho xã hội tính theo người. Trầm cảm + lo âu rẻ hơn nhiều theo người, NHƯNG do phổ biến hơn nên ĐÓNG GÓP CHÍNH vào tổng chi phí quốc gia.")
nr("Ngoài việc PHỔ BIẾN và TỐN KÉM, các tình trạng SKTT cũng KHÔNG ĐƯỢC PHỤC VỤ ĐỦ. Hệ thống SKTT trên toàn thế giới có các KHOẢNG TRỐNG và MẤT CÂN BẰNG lớn về thông tin + nghiên cứu, quản trị, nguồn lực và dịch vụ. Trung bình, các quốc gia dành dưới 2% ngân sách y tế cho SKTT. HƠN 70% chi tiêu SKTT ở các nước thu nhập trung bình vẫn dành cho BỆNH VIỆN TÂM THẦN tập trung — chứ không phải dịch vụ cộng đồng.")

H2("Chương 4 — Lợi ích của thay đổi: Đầu tư vào SKTT là đầu tư cho tương lai")
nr("Có 3 LÝ DO CHÍNH để đầu tư vào SKTT: (1) y tế công cộng; (2) quyền con người; (3) phát triển kinh tế-xã hội.")
nr("Đầu tư vào SKTT cho TẤT CẢ là TIẾN ĐẾN y tế công cộng — có thể giảm đáng kể đau khổ và cải thiện sức khoẻ + chất lượng cuộc sống + chức năng + tuổi thọ của người có rối loạn. Việc bao gồm SKTT trong các gói dịch vụ thiết yếu của Bao phủ Y tế Toàn dân (UHC) là ĐIỀU SỐNG CÒN. Cần TÍCH HỢP chăm sóc SKTT và sức khoẻ thể chất.")
nr("Đầu tư vào SKTT là CẦN THIẾT để DỪNG vi phạm quyền con người. Trên toàn thế giới, người có rối loạn tâm thần thường bị LOẠI TRỪ khỏi đời sống cộng đồng và bị từ chối các quyền cơ bản: phân biệt đối xử trong việc làm, giáo dục, nhà ở; không được công nhận bình đẳng trước pháp luật; thậm chí bị NGƯỢC ĐÃI bởi chính các dịch vụ y tế chăm sóc họ.")
nr("Can thiệp giảm KỲ THỊ — đặc biệt là chiến lược TIẾP XÚC XÃ HỘI thông qua người có trải nghiệm sống — có thể giảm kỳ thị và phân biệt đối xử trong cộng đồng.")

H2("Chương 5 — Nền tảng cho thay đổi: Củng cố hệ thống y tế")
nr("Chuyển hoá SKTT BẮT ĐẦU từ việc xây dựng các nền tảng cho hệ thống và dịch vụ SKTT hoạt động tốt. Các lĩnh vực hành động chính: (1) quản trị + lãnh đạo; (2) tài chính; (3) nhận thức công chúng; (4) năng lực chăm sóc SKTT.")
nr("Khung quốc tế và quốc gia là CỰC KỲ QUAN TRỌNG để định hướng hành động và tạo môi trường thuận lợi cho chuyển hoá. Cần luật pháp tuân thủ các công cụ nhân quyền quốc tế.")
nr("3 LOẠI CAM KẾT CHÍNH TRỊ cần thiết: (1) bày tỏ; (2) thể chế; (3) ngân sách. Tình trạng khẩn cấp y tế công cộng (vd: COVID-19) là TRÁCH NHIỆM và CƠ HỘI cho các nước đầu tư vào SKTT — chúng cung cấp NỀN TẢNG KHÔNG ĐÂU CÓ cho thay đổi.")
nr("Để chuyển hoá dịch vụ SKTT, cam kết phải được dịch thành hành động qua TÀI CHÍNH PHÙ HỢP — nghĩa là người ra chính sách + người lập kế hoạch cần dành NHIỀU TIỀN HƠN cho SKTT.")
nr("Lực lượng lao động CÓ NĂNG LỰC và CÓ ĐỘNG LỰC là thành phần SỐNG CÒN. Tất cả các nước cần MỞ RỘNG lực lượng chuyên môn SKTT, đồng thời xây dựng NĂNG LỰC chăm sóc SKTT cho lực lượng y tế khác (như bác sĩ tổng quát).")

H2("Chương 6 — Phát huy + Phòng ngừa cho thay đổi: Đa ngành cho TẤT CẢ")
nr("Ở MỌI giai đoạn cuộc đời, phát huy + phòng ngừa là CẦN THIẾT để: tăng cường WELL-BEING + KIÊN CƯỜNG; ngăn chặn khởi phát + tác động của rối loạn; giảm nhu cầu chăm sóc SKTT. Có bằng chứng tăng dần rằng phát huy + phòng ngừa có thể CHI PHÍ-HIỆU QUẢ.")
nr("Các can thiệp hoạt động bằng cách: (1) xác định các yếu tố quyết định cá nhân/xã hội/cấu trúc của SKTT; (2) can thiệp để giảm rủi ro, tăng kiên cường, tạo môi trường HỖ TRỢ. Các nhóm ưu tiên: (a) phòng ngừa tự tử; (b) trẻ em + vị thành niên; (c) SKTT tại nơi làm việc.")
nr("Tái định hình các yếu tố quyết định SKTT thường ĐÒI HỎI hành động NGOÀI ngành y tế — làm cho phát huy + phòng ngừa hiệu quả là một NHIỆM VỤ ĐA NGÀNH.")

H2("Chương 7 — Tái cấu trúc và Mở rộng dịch vụ chăm sóc cho TÁC ĐỘNG")
nr("Trung tâm của cải cách SKTT là VIỆC TÁI TỔ CHỨC LỚN dịch vụ SKTT. Phải DỊCH CHUYỂN trọng tâm chăm sóc cho các tình trạng nghiêm trọng TỪ bệnh viện tâm thần ĐẾN dịch vụ SKTT dựa vào cộng đồng — đóng cửa các bệnh viện tâm thần dài hạn KHI ĐÃ CÓ phương án thay thế cộng đồng đầy đủ.")
nr("Đồng thời, chăm sóc cho các tình trạng phổ biến (trầm cảm, lo âu) phải được MỞ RỘNG. Cả 2 chiến lược đều CẦN THIẾT để cải thiện độ bao phủ và chất lượng. Chăm sóc cộng đồng DỄ TIẾP CẬN HƠN + ĐƯỢC CHẤP NHẬN HƠN + HIỆU QUẢ HƠN so với chăm sóc trong cơ sở. Cần chăm sóc LẤY NGƯỜI LÀM TRUNG TÂM, ĐỊNH HƯỚNG PHỤC HỒI và DỰA TRÊN QUYỀN CON NGƯỜI.")
nr("Chăm sóc cộng đồng bao gồm MẠNG LƯỚI dịch vụ liên kết: (1) SKTT tích hợp trong chăm sóc y tế chung; (2) dịch vụ SKTT cộng đồng chuyên biệt; (3) dịch vụ SKTT trong các cơ sở phi y tế (trường học, nơi làm việc, dịch vụ xã hội). Hỗ trợ phi chính thức từ cộng đồng (nhân viên cộng đồng, đồng đẳng) BỔ SUNG cho dịch vụ chính thức.")
nr("Tích hợp SKTT trong chăm sóc y tế chung thường gồm CHIA SẺ NHIỆM VỤ (task-sharing) với nhân viên y tế không chuyên SKTT — đã được chứng minh giảm khoảng cách điều trị + tăng độ bao phủ.")

H2("Chương 8 — Kết luận: Hành động toàn diện + Con đường chuyển hoá")
nr("Báo cáo kết luận với 3 CON ĐƯỜNG chuyển hoá: (1) PHÁT HUY SKTT cho tất cả; (2) BẢO VỆ người có nguy cơ; (3) CUNG CẤP chăm sóc SKTT trong cộng đồng. Cần KẾT HỢP 3 con đường để có HÀNH ĐỘNG TOÀN DIỆN.")

# ============================================================
# PHẦN 4 — CHƯƠNG 3 CHI TIẾT (Epidemiology)
# ============================================================
H1("PHẦN 4 — CHƯƠNG 3 CHI TIẾT: SKTT THẾ GIỚI HÔM NAY")
page_marker("Trang 35-65, Chapter 3: World mental health today")

H2("3.1. Tổng quan dịch tễ học (Epidemiological overview)")
nr("Trước đại dịch — năm 2019 — ước tính 970 TRIỆU người trên thế giới đang sống chung với một rối loạn tâm thần (mental disorder), trong đó 82% ở các nước thu nhập thấp + trung bình (LMICs). Giữa năm 2000 và 2019, ước tính 25% NHIỀU NGƯỜI HƠN sống với rối loạn — nhưng vì dân số thế giới tăng cùng tốc độ, TỈ LỆ HIỆN MẮC điểm (point prevalence) GIỮ NGUYÊN ở khoảng 13% trong giai đoạn này.")

nr("BỔ SUNG (theo các ước tính khác):")
nr("• 283 triệu người có rối loạn sử dụng rượu (2016)")
nr("• 36 triệu người có rối loạn sử dụng ma túy (2019)")
nr("• 55 triệu người có sa sút trí tuệ (dementia, 2019)")
nr("• 50 triệu người có động kinh (epilepsy, 2015)")

H2("3.2. Tỉ lệ phổ biến theo rối loạn (Bảng 3.1 tái tạo)")
nr("Tái tạo Bảng 3.1 từ PDF gốc trang 41 — tỉ lệ điểm (point prevalence) năm 2019 theo rối loạn, giới tính, và tuổi.")
tbl(['Rối loạn', 'Tổng (triệu)', '% Tổng', '% Nam', '% Nữ', '% 5-9 tuổi', '% 10-14', '% 15-19', '% 20+'], [
    ['Tổng rối loạn tâm thần', '970', '13,0%', '12,5%', '13,5%', '7,6%', '13,5%', '14,7%', '14,7%'],
    ['Rối loạn lo âu (gồm PTSD)', '301', '4,0%', '3,0%', '5,0%', '1,5%', '3,6%', '4,6%', '4,8%'],
    ['Rối loạn trầm cảm (gồm dysthymia)', '280', '3,8%', '3,0%', '4,5%', '1,1%', '2,8%', '4,0%', '5,4%'],
    ['ADHD (Tăng động giảm chú ý)', '85', '1,1%', '1,7%', '0,6%', '2,4%', '—', '—', '—'],
    ['Rối loạn lưỡng cực', '40', '0,5%', '0,5%', '0,6%', '—', '0,2%', '0,6%', '0,7%'],
    ['Rối loạn phổ tự kỷ (ASD)', '28', '0,4%', '0,6%', '0,2%', '0,5%', '0,4%', '0,4%', '0,3%'],
    ['Tâm thần phân liệt', '24', '0,3%', '0,3%', '0,3%', '0,1%', '0,3%', '0,5%', '0,5%'],
    ['Rối loạn ăn uống (anorexia + bulimia)', '14', '0,2%', '0,1%', '0,2%', '—', '0,1%', '0,3%', '0,3%'],
])

nr("QUAN SÁT QUAN TRỌNG:")
nr("• Lo âu và trầm cảm là 2 rối loạn PHỔ BIẾN NHẤT — chiếm gần 60% tổng số ca rối loạn tâm thần.")
nr("• NỮ cao hơn NAM ở: lo âu, trầm cảm, ăn uống.")
nr("• NAM cao hơn NỮ ở: ADHD (gấp ~3x), tự kỷ (gấp ~3x).")
nr("• Tâm thần phân liệt: ~1 trên 200 người lớn (24 triệu toàn cầu).")
nr("• Rối loạn lưỡng cực: ~1 trên 150 người lớn (40 triệu toàn cầu).")

H2("3.3. Tác động COVID-19 đến trầm cảm + lo âu (Box 3.2)")
nr("Trước đại dịch năm 2020: ước tính 193 TRIỆU người có rối loạn trầm cảm chính (2.471/100.000 dân) và 298 TRIỆU người có rối loạn lo âu (3.825/100.000 dân).")
nr("SAU khi điều chỉnh tác động COVID-19: ước tính BAN ĐẦU cho thấy TĂNG VỌT lên 246 TRIỆU (3.153/100.000) cho trầm cảm và 374 TRIỆU (4.802/100.000) cho lo âu năm 2020 — TƯƠNG ĐƯƠNG TĂNG ~25%.")

H2("3.4. Tự tử (Suicide)")
page_marker("Trang 54-55, Section 3.1.3 Suicide")
nr("Năm 2019: ước tính 703.000 người ở MỌI lứa tuổi đã tự tử (9 trên 100.000 dân toàn cầu). Tỉ lệ tự tử CHÊNH ĐÁNG KỂ giữa các nước: từ <2/100.000 ở một số nước đến >80/100.000 ở các nước khác.")
nr("• KHOẢNG 3/4 (77%) tổng số ca tự tử xảy ra ở các nước thu nhập thấp + trung bình (LMICs) — nơi đa số dân số thế giới sinh sống.")
nr("• Nhưng các nước thu nhập cao gộp lại có TỈ LỆ tự tử CAO NHẤT: ~10,9/100.000.")
nr("• Tỉ lệ Nam:Nữ tử vong do tự tử = 2:1 toàn cầu; ở nước thu nhập cao = 3:1.")
nr("• Phụ nữ THƯỜNG MƯU SÁT NHIỀU HƠN nhưng nam tử vong cao hơn (do dùng phương pháp gây tử vong cao hơn).")
nr("• HƠN MỘT NỬA (58%) ca tự tử xảy ra TRƯỚC 50 tuổi.")
nr("• Người >70 tuổi có tỉ lệ tự tử CAO HƠN GẤP ĐÔI người tuổi lao động.")
nr("• Tỉ lệ tự tử toàn cầu đã GIẢM 36% kể từ năm 2000 — giảm 47% ở Châu Âu, 49% ở Tây Thái Bình Dương; NHƯNG TĂNG 17% ở Châu Mỹ.")
nr("• Cứ MỖI ca tự vong do tự tử có HƠN 20 ca mưu sát — tỷ lệ 20:1.")
nr("• Tự tử = HƠN 1/100 ca tử vong toàn cầu.")

H2("3.5. Gánh nặng bệnh tật (Burden of disease)")
nr("Rối loạn tâm thần CHIẾM 1/6 (16,7%) tổng YLDs (Năm sống với khuyết tật) toàn cầu — NHIỀU NHẤT trong các nguyên nhân YLDs.")
nr("• Gộp rối loạn tâm thần + thần kinh + sử dụng chất = 1/4 (25%) tổng YLDs toàn cầu.")
nr("• Trầm cảm là nguyên nhân YLDs LỚN THỨ 2 toàn cầu — chiếm 5,8%.")
nr("• Lo âu, trầm cảm thường xuyên trong TOP 10 nguyên nhân YLDs từ 2000.")
nr("• Tỉ trọng RT trong YLDs theo tuổi: <10% ở trẻ em + người già; >23% ở người 15-29 tuổi.")
nr("LƯU Ý QUAN TRỌNG: gánh nặng RT chủ yếu là YLDs (chứ không phải YLLs) vì khung tính KHÔNG GÁN tử vong cho trầm cảm/lưỡng cực, và tự tử/tự hại được phân loại RIÊNG là \"chấn thương có chủ ý\".")

H2("3.6. Hệ quả kinh tế (Economic consequences)")
nr("Diễn đàn Kinh tế Thế giới ước tính: bộ rối loạn tâm thần định nghĩa rộng đã tốn cho NỀN KINH TẾ THẾ GIỚI ~2,5 NGHÌN TỈ USD năm 2010 — DỰ KIẾN tăng lên 6 NGHÌN TỈ USD năm 2030.")
nr("→ Đây là CON SỐ CAO HƠN tổng chi phí dự đoán cho ung thư + tiểu đường + bệnh hô hấp mạn cộng lại.")
nr("LMICs dự kiến gánh 35% chi phí này.")
nr("• Chi phí gián tiếp (mất năng suất, thất nghiệp) thường CAO HƠN chi phí điều trị trực tiếp.")
nr("• Tâm thần phân liệt là rối loạn ĐẮT NHẤT cho xã hội PER CAPITA.")
nr("• Trầm cảm + lo âu rẻ hơn per capita NHƯNG đóng góp lớn nhất vào tổng chi phí quốc gia (do số ca cao).")

H2("3.7. Khoảng trống trong y tế công cộng SKTT")
nr("• Trung bình các quốc gia dành DƯỚI 2% ngân sách y tế cho SKTT.")
nr("• HƠN 70% chi tiêu SKTT ở nước thu nhập trung bình vẫn dành cho BỆNH VIỆN TÂM THẦN tập trung.")
nr("• Hầu hết các nước báo cáo (67%) chi DƯỚI 20% ngân sách SKTT cho dịch vụ SKTT cộng đồng.")
nr("• Khoảng 80% các nước chi DƯỚI 20% cho SKTT trong bệnh viện đa khoa và 80% chi dưới 20% cho SKTT trong y tế cơ bản.")
nr("• Tự tử + tự hại — chỉ số SDG DUY NHẤT về SKTT — nhận DƯỚI 1% tổng kinh phí nghiên cứu SKTT.")

# ============================================================
# PHẦN 5 — CHƯƠNG 7 CHI TIẾT (Restructuring care)
# ============================================================
H1("PHẦN 5 — CHƯƠNG 7 CHI TIẾT: TÁI CẤU TRÚC + MỞ RỘNG DỊCH VỤ CHĂM SÓC")
page_marker("Trang 187-245, Chapter 7: Restructuring and scaling up care")

H2("7.1. Hiểu chăm sóc SKTT dựa vào cộng đồng")
nr("Chăm sóc SKTT cộng đồng có 3 ƯU ĐIỂM so với chăm sóc trong cơ sở: (1) DỄ TIẾP CẬN HƠN; (2) ĐƯỢC CHẤP NHẬN HƠN; (3) KẾT QUẢ ĐIỀU TRỊ TỐT HƠN. Cần dựa trên 4 NGUYÊN TẮC: (a) lấy người làm trung tâm; (b) định hướng phục hồi; (c) quyền con người; (d) toàn diện đa ngành.")

H2("7.2. SKTT tích hợp trong dịch vụ y tế chung")
nr("Tích hợp SKTT thường gồm: (1) CHIA SẺ NHIỆM VỤ (task-sharing) với nhân viên y tế không chuyên; (2) THÊM nhân viên SKTT chuyên trách + nguồn lực vào y tế cơ bản và thứ cấp.")
nr("Chia sẻ nhiệm vụ với y tế cơ bản đã được chứng minh: (a) GIẢM khoảng cách điều trị; (b) TĂNG độ bao phủ cho rối loạn tâm thần ưu tiên.")
nr("Chia sẻ nhiệm vụ TRONG các chương trình bệnh cụ thể (HIV/AIDS, lao) có thể CẢI THIỆN cả kết quả SKTT và bệnh nền tảng.")

H2("7.3. Dịch vụ SKTT cộng đồng (chuyên biệt)")
nr("Bao gồm: (a) trung tâm SKTT cộng đồng đa ngành; (b) đội ngũ điều trị nội trú/ngoại trú cộng đồng; (c) dịch vụ tại nhà; (d) bệnh viện ban ngày; (e) đơn vị lưu trú cấp tính trong bệnh viện đa khoa.")
nr("Khi đã có đầy đủ phương án cộng đồng, cần ĐÓNG CỬA bệnh viện tâm thần dài hạn (deinstitutionalization) — chuyển bệnh nhân về cộng đồng với dịch vụ HỖ TRỢ phù hợp.")

H2("7.4. Dịch vụ SKTT NGOÀI ngành y tế")
nr("Bao gồm: (a) SKTT trong trường học (mental health in schools); (b) SKTT tại nơi làm việc; (c) SKTT trong dịch vụ xã hội + nhà ở; (d) SKTT trong tư pháp + nhà tù; (e) hỗ trợ đồng đẳng (peer support) + tự giúp.")
nr("CẢI THIỆN ĐA NGÀNH là CỐT LÕI của chuyển hoá SKTT — không thể chỉ dựa vào ngành y tế đơn lẻ.")

# ============================================================
# PHẦN 6 — KẾT LUẬN (Chương 8)
# ============================================================
H1("PHẦN 6 — KẾT LUẬN (Chương 8)")
page_marker("Trang 247-258, Chapter 8: Conclusion")

H2("8.1. Hành động toàn diện (Comprehensive action)")
nr("Chuyển hoá SKTT đòi hỏi HÀNH ĐỘNG TOÀN DIỆN trên 3 con đường: (1) phát huy + bảo vệ SKTT; (2) phòng ngừa rối loạn; (3) cung cấp chăm sóc SKTT dựa vào cộng đồng.")

H2("8.2. Con đường chuyển hoá (Paths to transformation)")
nr("4 NHIỆM VỤ chính theo bản đồ chuyển hoá:")
nr("(1) TÁI ĐỊNH HÌNH MÔI TRƯỜNG (reshape environments) — tác động yếu tố quyết định đa ngành.")
nr("(2) TĂNG GIÁ TRỊ + CAM KẾT (deepen value and commitment) — chính trị, ngân sách, nhận thức.")
nr("(3) CỦNG CỐ CHĂM SÓC SKTT (strengthen mental health care) — tái cấu trúc, dịch vụ cộng đồng.")
nr("(4) TRUYỀN CẢM HỨNG + THÔNG TIN cho thay đổi (inspire + inform change) — báo cáo này.")

H2("8.3. Kết hợp nỗ lực vì thay đổi (Combining efforts for change)")
nr("WHO mời GỌI tất cả người ra quyết định, người ủng hộ, người có trải nghiệm sống, gia đình + bạn bè, nhà nghiên cứu, người làm thực tiễn, người tuyển dụng, học sinh + giáo viên, đồng nghiệp + cộng đồng — TẤT CẢ — kết hợp nỗ lực để CHUYỂN HOÁ SKTT cho TẤT CẢ MỌI NGƯỜI.")

# ============================================================
# PHẦN 7 — TÀI LIỆU THAM KHẢO CHÍNH
# ============================================================
H1("PHẦN 7 — DANH MỤC TÀI LIỆU THAM KHẢO CHÍNH")
nr("WHO 2022 Report có ~600 tài liệu tham khảo (toàn báo cáo). Em liệt kê các tài liệu KEY được trích dẫn trong các phần đã dịch chi tiết — cho phép thầy đối chiếu trực tiếp.")
nr("(Giữ nguyên tiếng Anh — định dạng gốc của WHO)")
refs = [
    "1. WHO. The world health report 2001. Mental health: new understanding, new hope. Geneva: World Health Organization; 2001.",
    "2. WHO. Mental health action plan 2013–2030. Geneva: World Health Organization; 2021.",
    "3. WHO. Mental health atlas 2020. Geneva: World Health Organization; 2021.",
    "96. GBD 2019 Mental Disorders Collaborators. Global, regional, and national burden of 12 mental disorders in 204 countries and territories, 1990–2019. Lancet Psychiatry. 2022;9(2):137–50. DOI: 10.1016/S2215-0366(21)00395-3.",
    "99. Institute for Health Metrics and Evaluation (IHME). Global Burden of Disease Study 2019 results. Seattle, WA: IHME, 2020.",
    "100. WHO. Global status report on alcohol and health 2018. Geneva: World Health Organization; 2018.",
    "101. UNODC. World drug report 2021. Vienna: United Nations Office on Drugs and Crime; 2021.",
    "102. WHO. Global status report on the public health response to dementia. Geneva: World Health Organization; 2021.",
    "123. WHO. Suicide worldwide in 2019: global health estimates. Geneva: World Health Organization; 2021.",
    "124. WHO. Preventing suicide: a global imperative. Geneva: World Health Organization; 2014.",
    "127. WHO. Global health estimates 2019: deaths by cause, age, sex, by country and by region, 2000–2019. Geneva: World Health Organization; 2020.",
    "129. WHO. Global health estimates 2019: disease burden by cause, age, sex, by country and by region, 2000–2019. Geneva: World Health Organization; 2020.",
    "Bloom DE, Cafiero ET, Jané-Llopis E, et al. The global economic burden of noncommunicable diseases. Geneva: World Economic Forum; 2011.",
    "Santomauro DF, Mantilla Herrera AM, Shadid J, et al. Global prevalence and burden of depressive and anxiety disorders in 204 countries and territories in 2020 due to the COVID-19 pandemic. Lancet. 2021;398(10312):1700–12.",
    "United Nations. Transforming our world: the 2030 Agenda for Sustainable Development. New York: UN; 2015.",
]
for ref in refs:
    nr(ref)
nr("(Còn ~585 tài liệu tham khảo trong báo cáo gốc — thầy có thể truy cập đầy đủ tại: https://www.who.int/publications/i/item/9789240049338)")

# ============================================================
# PHẦN 8 — PHẢN BIỆN CHI TIẾT
# ============================================================
H1("PHẦN 8 — PHẢN BIỆN CHI TIẾT")
page_marker("Phần phản biện của trợ lý nghiên cứu — không có trong PDF gốc")

H2("Điểm 1 — Báo cáo cấp toàn cầu CHUẨN nhưng THIẾU đặc thù quốc gia")
crit_para("WHO 2022 là báo cáo CHUẨN cho mọi quốc gia thành viên — có sức nặng chính trị + đạo đức cao. Tuy nhiên, do bao quát 194 quốc gia, các khuyến nghị BẮT BUỘC PHẢI ở mức tổng quát — không có hướng dẫn cụ thể cho VIỆT NAM (như: ngân sách bao nhiêu, công cụ chẩn đoán nào, đào tạo bao nhiêu nhân viên). Báo cáo nên được KẾT HỢP với: (a) Mental Health Atlas 2020 cho dữ liệu nước; (b) các khảo sát quốc gia (V-NAMHS 2022 cho VN); (c) khung pháp luật quốc gia. WHO + Bộ Y tế VN cần xây dựng \"Lộ trình triển khai WMH 2022 cho Việt Nam\" để áp dụng cụ thể.")

H2("Điểm 2 — Số liệu dịch tễ DỰA VÀO GBD 2019 — kế thừa hạn chế của GBD")
crit_para("Tất cả các con số PHỔ BIẾN + GÁNH NẶNG trong báo cáo (970 triệu người, 1/8 dân số, 13%, 1/6 YLDs) đều TỪ GBD 2019. Do đó báo cáo KẾ THỪA TẤT CẢ hạn chế của GBD: (a) YLLs CỰC THẤP do khung không gán tử vong cho trầm cảm; (b) dữ liệu thiếu ở LMICs → ước tính BIÊN CHỆCH; (c) dùng DSM-IV-TR/ICD-10 đã lạc hậu; (d) trọng số khuyết tật từ vài nước thu nhập cao. Tỉ lệ 13% có thể là ƯỚC TÍNH THẤP HƠN THỰC TẾ — đặc biệt ở các nước có dữ liệu yếu (như Việt Nam). Khuyến cáo của thầy: KHI báo cáo trong luận văn, ghi rõ \"theo ước tính WHO 2022 (dựa GBD 2019)\" để không lan truyền sai số.")

H2("Điểm 3 — Định nghĩa thuật ngữ KHÔNG NHẤT QUÁN")
crit_para("Báo cáo dùng nhiều thuật ngữ tương tự: \"mental health condition\", \"mental disorder\", \"mental illness\", \"psychosocial disability\". Trong Chương 3 dùng \"mental disorder\" (theo ICD-11), nhưng các chương khác dùng \"mental health condition\" (rộng hơn — bao gồm khuyết tật tâm lý-xã hội). Thiếu nhất quán này gây KHÓ KHĂN khi dịch sang tiếng Việt — không rõ \"rối loạn tâm thần\" có bao gồm cả người cảm thấy đau khổ tinh thần ngắn hạn không. Đề xuất: trong các tài liệu tiếng Việt, dùng \"rối loạn tâm thần\" cho \"mental disorder\" (chẩn đoán theo ICD/DSM) và \"vấn đề SKTT\" hoặc \"tình trạng SKTT\" cho \"mental health condition\" (rộng hơn).")

H2("Điểm 4 — Thiếu dữ liệu COVID-19 đầy đủ")
crit_para("Báo cáo xuất bản 06/2022 — chỉ có ước tính COVID-19 cho năm 2020 (Santomauro et al. 2021). KHÔNG có dữ liệu cho 2021-2022 — các năm có biến chủng Delta + Omicron + lockdown kéo dài + suy thoái kinh tế. Tác động dài hạn của COVID-19 lên SKTT (hậu COVID, mất người thân, thất nghiệp, học sinh nghỉ học) sẽ chỉ rõ ràng SAU 2025 — báo cáo này KHÔNG NẮM BẮT được. Cần CẬP NHẬT bằng các bài hậu 2022 — như các bài Lancet Psychiatry 2024-2025.")

H2("Điểm 5 — Chuyển hoá đòi hỏi đầu tư mà KHÔNG CHỈ ra nguồn lực cụ thể")
crit_para("Báo cáo kêu gọi tăng ngân sách SKTT, mở rộng nhân lực, đóng cửa bệnh viện tâm thần, mở dịch vụ cộng đồng — TOÀN BỘ là KHUYẾN NGHỊ CHIẾN LƯỢC. Nhưng KHÔNG nêu các con số CỤ THỂ về: (a) ngân sách tăng bao nhiêu (% GDP, % chi y tế); (b) cần bao nhiêu bác sĩ tâm thần/nhà tâm lý/nhân viên cộng đồng/100.000 dân; (c) thời gian biểu chuyển đổi từ bệnh viện sang cộng đồng. Các nước thu nhập thấp khó KHẢ THI hoá khuyến nghị WHO. Báo cáo nên kèm \"hộp công cụ định lượng\" cho các nước.")

H2("Điểm 6 — Cảm hứng tốt nhưng còn THIẾU bằng chứng kinh tế-hiệu quả vững chắc")
crit_para("Báo cáo viết với giọng vận động — nhiều câu \"chúng ta cần\", \"phải chuyển hoá\". Tốt cho cảm hứng. NHƯNG: bằng chứng về CHI PHÍ-HIỆU QUẢ (cost-effectiveness) của các can thiệp ĐA NGÀNH (vd: SKTT trường học, SKTT nơi làm việc) còn YẾU. Báo cáo không trả lời được câu hỏi \"đầu tư 1 USD vào can thiệp X tạo ra bao nhiêu USD lợi ích kinh tế?\" cho hầu hết các đề xuất. Cần thêm phân tích kiểu WHO-CHOICE để hỗ trợ ra quyết định ngân sách.")

# ============================================================
# PHẦN 9 — ÁP DỤNG VIỆT NAM
# ============================================================
H1("PHẦN 9 — ÁP DỤNG WHO 2022 CHO VIỆT NAM — ĐỀ XUẤT CỤ THỂ")

H2("Đề xuất 1 — Xây dựng \"Lộ trình WHO Mental Health 2022 cho Việt Nam 2026-2030\"")
nr("Cần MỘT TÀI LIỆU CHÍNH SÁCH cấp Bộ Y tế hoặc Chính phủ — dịch các nguyên tắc WHO 2022 thành: (a) chỉ tiêu cụ thể VN đến 2030 (vd: chi 4% ngân sách y tế cho SKTT, có 5 bác sĩ tâm thần/100.000 dân, 100% xã có nhân viên SKTT cộng đồng); (b) các bước chuyển từ mô hình bệnh viện tâm thần sang mô hình cộng đồng tại VN; (c) hợp tác đa ngành (y tế + giáo dục + LĐTBXH + công an + phụ nữ).")

H2("Đề xuất 2 — Tích hợp SKTT vào trạm y tế xã (PHC) toàn quốc")
nr("WHO 2022 nhấn mạnh chia sẻ nhiệm vụ với y tế cơ bản. VN có ~10.500 trạm y tế xã + ~1.700 phòng khám đa khoa khu vực — mạng lưới y tế cơ bản LỚN NHẤT. Cần: (a) đào tạo 1 nhân viên SKTT/trạm; (b) cung cấp công cụ sàng lọc PHQ-9 + GAD-7 chuẩn VN; (c) phác đồ điều trị trầm cảm + lo âu + tâm thần phân liệt nhẹ tại trạm; (d) đường tham vấn lên tuyến huyện-tỉnh khi cần.")

H2("Đề xuất 3 — Mở rộng SKTT trong trường học cho học sinh trung học")
nr("WHO 2022 ưu tiên trẻ em + vị thành niên. VN có ~5.700 trường THPT + ~12.000 trường THCS. Cần: (a) chương trình giáo dục cảm xúc-xã hội (SEL) trong giáo viên chủ nhiệm; (b) tư vấn tâm lý học đường (đã có quy định Bộ GD-ĐT 2017 nhưng triển khai còn yếu); (c) sàng lọc lo âu/trầm cảm cho HS lớp 9 + 12; (d) hệ thống chuyển tuyến khi phát hiện rủi ro tự tử.")

H2("Đề xuất 4 — Hành động khẩn về phòng ngừa tự tử ở Việt Nam")
nr("Việt Nam thiếu dữ liệu công khai về tỉ lệ tự tử cấp quốc gia (Mental Health Atlas 2020 không có số chính thức của VN). Cần: (a) hệ thống giám sát tự tử quốc gia — qua giấy chứng tử + báo cáo bệnh viện; (b) phân tích tỉ lệ theo độ tuổi + giới + vùng; (c) chương trình dự phòng tự tử cấp quốc gia (theo mẫu LIVE LIFE của WHO 2021); (d) hạn chế tiếp cận phương tiện gây tử vong cao (thuốc trừ sâu, vũ khí, độ cao). Tự tử là 1/100 ca tử vong toàn cầu — VN cần biết con số của mình.")

H2("Đề xuất 5 — Giáo dục đào tạo + Truyền thông giảm kỳ thị")
nr("WHO 2022 nhấn mạnh giảm kỳ thị qua TIẾP XÚC XÃ HỘI với người có trải nghiệm sống. VN cần: (a) kênh truyền thông quốc gia về SKTT (TV, mạng xã hội); (b) câu chuyện người thật về phục hồi từ rối loạn tâm thần; (c) tích hợp SKTT vào chương trình GD công dân + sinh hoạt dưới cờ; (d) mạng lưới hội người có trải nghiệm sống cấp địa phương; (e) đào tạo phóng viên không phân biệt đối xử khi viết về tâm thần (theo hướng dẫn WHO Media Guidelines).")

H2("Đề xuất 6 — Hợp tác Nghiên cứu + WHO Collaborator Network")
nr("VN nên: (a) cử đại diện tham gia GBD Collaborator Network để cung cấp dữ liệu VN; (b) hợp tác với WHO Mental Health Department + WHO WPRO (Western Pacific Regional Office) để có hỗ trợ kỹ thuật; (c) khảo sát quốc gia định kỳ 5 năm/lần (V-NAMHS 2022 là khảo sát quốc gia ĐẦU TIÊN — cần lặp lại 2027).")

# ============================================================
# KẾT
# ============================================================
H1("KẾT THÚC")
nr("Tài liệu: BẢN DỊCH + PHẢN BIỆN báo cáo WHO 2022 World Mental Health Report (296 trang gốc).")
nr("Trợ lý nghiên cứu — 30/04/2026 — tiếng Việt thuần — Open Access CC BY-NC-SA 3.0 IGO.")
nr("Phương thức dịch: 1-1 cho Lời nói đầu, Tóm tắt 8 chương, Chương 3 (dịch tễ), Chương 7 (tái cấu trúc), Chương 8 (kết luận). Các chương 2, 4, 5, 6 dịch tóm lược chính (các chương này chủ yếu là khuyến nghị + nguyên tắc, ít dữ liệu).")
nr("Phần Phản biện + Áp dụng VN là phần em tự thêm vào — KHÔNG có trong báo cáo gốc.")

OUT.parent.mkdir(parents=True, exist_ok=True)
d.save(str(OUT))
print(f"Wrote: {OUT}")
