"""DOC: Tong hop AUDIT TOAN DIEN cho 2 doc moi (H4 + BNHD)
- 5 vong kiem tra
- Cross-check RAG (102 entries) + KG (184 -> 208 nodes)
- Gap KG da sua: QT009 missing + 11 concepts mới + 5 papers ngoài + 3 docs + 2 questions
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/AUDIT_TOAN_DIEN_RAG_KG_H4_BNHD.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('AUDIT TOÀN DIỆN — 2 DOC MỚI (H4 GENDER + BNHĐ)\nKIỂM TRA NỘI DUNG + RAG + KG\n— 5 vòng + cross-check 4 nguồn dữ liệu (08/05/2026) —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Tổng kết
H('Tổng kết — KHÔNG còn lỗi nội dung; ĐÃ SỬA 1 GAP KG', level=2, color=GREEN)
para(
    'Sau audit toàn diện 4 cấp (5 vòng kiểm tra nội dung + cross-check '
    'canonical + RAG + KG), kết luận:', size=12, bold=True
)
para(
    '• 2 doc nội dung: SẠCH — 104 số liệu + 8 reference đều verified, không '
    'có fact bịa đặt.', indent=True
)
para(
    '• Canonical: 4/4 paper IDs (QT007, QT008, QT009, QT067) khớp 100%.', indent=True
)
para(
    '• RAG main index: 4/4 paper entries có đầy đủ.', indent=True
)
para(
    '• KG: PHÁT HIỆN 1 GAP — QT009 (Qiu 2022) thiếu trong KG paper nodes '
    'dù được cited trong Doc 4. Đã bổ sung + 5 paper external + 12 '
    'concept + 2 topic + 3 doc + 2 question + 52 edge.', indent=True
)

# I. Nội dung 2 doc
H('I. Nội dung 2 doc — 5 vòng kiểm tra', level=2, color=NAVY)
add_table(
    ['Vòng', 'Nội dung kiểm tra', 'Kết quả'],
    [
        ['1', 'Audit 104 số liệu (62 H4 + 42 BNHĐ)', '0 lỗi'],
        ['2', 'Verify references: Nolen-Hoeksema 2012 + OBPP', '0 lỗi'],
        ['3', 'Audit math: 1/0,262=3,82; 17,5/11,1=1,58; 17,5−11,1=6,4', '0 lỗi'],
        ['4', 'Cross-check IDs với canonical', '0 lỗi (4/4 ID khớp)'],
        ['5', 'Audit cuối cùng', '0 lỗi'],
    ]
)

# II. Cross-check 4 nguồn
H('II. Cross-check 4 nguồn dữ liệu', level=2, color=NAVY)
add_table(
    ['Nguồn', 'Kiểm tra', 'Kết quả'],
    [
        ['canonical_index.json (91 entries)', 'QT007, QT008, QT009, QT067', '✓ 4/4 khớp'],
        ['papers_metadata.json (91 entries)', 'Year + descriptor + region', '✓ 4/4 khớp'],
        ['rag_main_index.json (102 entries)', '4 paper IDs có entries với tokens', '✓ 4/4 có đầy đủ'],
        ['questions_kg.json (184 nodes ban đầu)', 'Paper nodes + concepts', '⚠ QT009 THIẾU + 11 concepts thiếu'],
    ]
)

# III. Gap KG đã sửa
H('III. Gap KG đã phát hiện và bổ sung', level=2, color=NAVY)
para('Trước update KG: 184 nodes / 381 edges', bold=True)
para('Sau update KG: 208 nodes / 433 edges (+24 nodes, +52 edges)', bold=True)
para('')

H('III.1. 2 Topic mới', level=3, color=NAVY)
add_table(
    ['ID', 'Label'],
    [
        ['TOPIC_GENDER_PSYCHOPATHOLOGY', 'Gender differences in psychopathology'],
        ['TOPIC_SCHOOL_BULLYING', 'School bullying'],
    ]
)

H('III.2. 12 Concept mới', level=3, color=NAVY)
add_table(
    ['ID', 'Label', 'Liên quan đến'],
    [
        ['CONCEPT_GENDER_DIFFERENCES', 'Gender differences in anxiety', 'Doc 4'],
        ['CONCEPT_SEX_RATIO', 'Sex ratio (female:male)', 'Doc 4'],
        ['CONCEPT_HPG_AXIS', 'HPG axis', 'Doc 4 cơ chế'],
        ['CONCEPT_RUMINATION', 'Rumination (Nolen-Hoeksema)', 'Doc 4 cơ chế'],
        ['CONCEPT_SCHOOL_BULLYING', 'School bullying', 'Doc BNHĐ'],
        ['CONCEPT_OBVQ', 'Olweus Bully/Victim Questionnaire', 'Doc BNHĐ'],
        ['CONCEPT_SCHOOL_REFUSAL', 'School refusal', 'Doc BNHĐ cơ chế'],
        ['CONCEPT_OBPP', 'Olweus Bullying Prevention Program', 'Doc BNHĐ hàm ý'],
        ['CONCEPT_SEPARATION_ANXIETY', 'RLLAC chia ly', 'Doc 4 + BNHĐ'],
        ['CONCEPT_SOCIAL_ANXIETY_DISORDER', 'RLLAXH xã hội', 'Doc 4 + BNHĐ'],
        ['CONCEPT_GENERALIZED_ANXIETY_DISORDER', 'RLLATQ lan tỏa', 'Doc 4 + BNHĐ'],
        ['CONCEPT_RCADS', 'Revised Children Anxiety Depression Scale', 'Cả 2 doc — thang đo chương 3'],
    ]
)

H('III.3. 5 Paper mới (1 chính + 4 external)', level=3, color=NAVY)
add_table(
    ['ID', 'Label', 'Lý do bổ sung'],
    [
        ['PAPER_QT009_Qiu_2022', 'QT009 Qiu (2022) Frontiers Public Health', '🔴 GAP — đã trong canonical/RAG nhưng THIẾU KG'],
        ['PAPER_Salk_2017_Psych_Bulletin', 'Salk, Hyde & Abramson (2017)', 'External quan trọng — Doc 4'],
        ['PAPER_McLean_2011_J_Psychiatr_Res', 'McLean et al. (2011)', 'External quan trọng — Doc 4'],
        ['PAPER_Olweus_1996_OBVQ', 'Olweus (1996) OBVQ-R', 'External quan trọng — Doc BNHĐ'],
        ['PAPER_NolenHoeksema_2012_Annu_Rev', 'Nolen-Hoeksema (2012)', 'External — Doc 4 cơ chế'],
    ]
)

H('III.4. 3 Doc + 2 Question + 52 Edge', level=3, color=NAVY)
add_table(
    ['ID', 'Loại', 'Mô tả'],
    [
        ['DOC_H4_GENDER_LOAI_RLLA', 'Doc', 'Gia_thuyet_khac_biet_gioi_tinh_theo_loai_RLLA.docx'],
        ['DOC_BNHD_BAT_NAT_HOC_DUONG', 'Doc', 'Dien_giai_yeu_to_bat_nat_hoc_duong.docx'],
        ['DOC_AUDIT_5_VONG_H4_BNHD', 'Doc', 'AUDIT_5_vong_H4_gender_va_BNHD.docx'],
        ['QA_23', 'Question', 'Phát biểu giả thuyết H4 gender × loại RLLA'],
        ['QA_24', 'Question', 'Diễn giải yếu tố BNHĐ tác động đến RLLA'],
        ['52 edges', 'Edge', 'BELONGS_TO + EXPLAINS + CITES + ANSWERED_IN + RELATED_TO'],
    ]
)

# IV. Trạng thái KG/RAG cuối
H('IV. Trạng thái KG/RAG sau update', level=2, color=NAVY)
add_table(
    ['Chỉ số', 'Trước', 'Sau', 'Chênh'],
    [
        ['KG nodes', '184', '208', '+24'],
        ['KG edges', '381', '433', '+52'],
        ['Paper nodes', '23', '28', '+5'],
        ['Concept nodes', '~105', '~117', '+12'],
        ['Topic nodes', '14', '16', '+2'],
        ['Question nodes', '22', '24', '+2'],
        ['Doc nodes', '20', '23', '+3'],
        ['RAG questions', '22', '24', '+2 (QA_23, QA_24)'],
    ]
)
para('')
para(
    'Đã sync cả 2 chatbot — light + heavy. File data đồng bộ.',
    bold=True
)

# V. Câu trả lời
H('V. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
p = d.add_paragraph()
r = p.add_run(
    'Sau audit toàn diện 4 cấp (5 vòng kiểm tra nội dung + cross-check '
    'canonical + RAG + KG): (1) 2 doc nội dung SẠCH — không có fact bịa '
    'đặt, 104 số liệu + 8 reference đều verified; (2) Canonical + RAG: '
    '4/4 paper IDs (QT007, QT008, QT009, QT067) khớp 100%; (3) KG '
    'PHÁT HIỆN 1 GAP: QT009 (Qiu 2022) thiếu node dù được cited; (4) '
    'Đã bổ sung KG: 24 nodes (1 paper QT009 + 4 paper external + 12 '
    'concept + 2 topic + 3 doc + 2 question) và 52 edges. Trạng thái '
    'cuối KG: 208 nodes / 433 edges; RAG questions: 24 entries. Cả 2 '
    'chatbot light + heavy đã sync.'
)
r.font.size = Pt(12); r.font.color.rgb = BLUE; r.italic = True

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
