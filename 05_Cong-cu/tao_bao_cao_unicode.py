# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_b(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def add_p(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_i(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def tbl(headers, rows):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    return t

# ==================== TITLE ====================
title = doc.add_heading('', level=0)
run = title.add_run('BÁO CÁO TỔNG HỢP\n3 BÀI NGHIÊN CỨU VIỆT NAM VỀ RỐI LOẠN LO ÂU')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_i('Ngày báo cáo: 25/03/2026. Nguồn: 3 bài PDF trong thư mục papers/Viet-nam/')
doc.add_paragraph()

# ==================== I. TÓM TẮT ====================
add_h('I. TÓM TẮT TỪNG BÀI', 1)

# --- BÀI 1 ---
add_h('Bài 1: Nguyễn Thị Vân (2020) — Yếu tố ảnh hưởng đến RLLA ở HS THPT tại TP.HCM', 2)

tbl(
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Nguyễn Thị Vân — ĐH KHXH&NV, ĐHQG TP.HCM'],
        ['Tạp chí', 'Tạp chí Khoa học Quản lý Giáo dục, Số 01(25), 3/2020'],
        ['ISSN', '2354–0788'],
        ['Mẫu', 'Sàng lọc 558 HS THPT; phân tích sâu 90 HS có RLLA'],
        ['Công cụ', 'Thang đo STAI (Spielberger), thích nghi hóa bởi Nguyễn Công Khanh (2000)'],
        ['Địa điểm', '4 trường THPT tại TP.HCM (nội thành + ngoại thành)'],
    ]
)
doc.add_paragraph()

add_b('Phương pháp:')
add_p('Nghiên cứu sử dụng thiết kế cắt ngang, sàng lọc 558 HS THPT bằng thang đo lo âu STAI của Spielberger, phiên bản đã được Nguyễn Công Khanh thích nghi hóa tại Việt Nam từ năm 2000. Kết quả sàng lọc xác định 15–18.5% HS có biểu hiện RLLA. Từ nhóm này, 90 HS được đưa vào phân tích sâu về các yếu tố ảnh hưởng, sử dụng mẫu thuận tiện tại 4 trường thuộc 2 khu vực (nội thành và ngoại thành). Dữ liệu được phân tích bằng tương quan Pearson và hồi quy tuyến tính đa biến (Stepwise) trên SPSS.')

add_b('Kết quả chính:')
add_p('Nghiên cứu xác định 4 nhóm yếu tố ảnh hưởng đến RLLA: (1) học tập, (2) gia đình, (3) quan hệ xã hội, (4) bản thân học sinh. Nhóm yếu tố "bản thân học sinh" có tương quan mạnh nhất với thang đo lo âu SAS (r=0.42), giải thích 85.4% sự biến thiên điểm lo âu. Ba yếu tố bản thân quan trọng nhất là: ít giao tiếp chia sẻ (X̄=0.91), cảm thấy cuộc sống thất vọng (X̄=0.90), và lo sợ khi gặp khó khăn (X̄=0.88). Điều này cho thấy HS thiếu kỹ năng ứng phó là yếu tố nội tại trực tiếp nhất dẫn đến RLLA.')

add_p('Nhóm yếu tố học tập chi phối mạnh thứ hai (r=0.37, giải thích 74%). Áp lực thi đại học chiếm tỷ lệ TX+RTX cao nhất (56.7%), tiếp theo là định hướng nghề nghiệp (51.5%) và áp lực kết quả thi (50.6%). Áp lực kỳ vọng từ cha mẹ về học tập cũng chiếm 48.9%. Yếu tố gia đình nổi bật: cha mẹ so sánh con với người khác (25.8%), kỳ vọng quá cao/thấp (25.1%), áp đặt la mắng (21%). Yếu tố xã hội: người xung quanh gây khó chịu (18.2%), thầy cô mang cảm xúc tiêu cực vào lớp (16%).')

doc.add_paragraph()

# --- BÀI 2 ---
add_h('Bài 2: Trần Thị Mỵ Lương (2020) — Rối loạn lo âu ở HS THPT Chuyên', 2)

tbl(
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Trần Thị Mỵ Lương & Đặng Đức Anh — Học viện Phụ nữ Việt Nam'],
        ['Tạp chí', 'Tạp chí Khoa học, Trường ĐH Thủ Đô Hà Nội, Số 40/2020'],
        ['Mẫu', '540 HS THPT Chuyên (3 khối 10–12), chọn mẫu ngẫu nhiên phân tầng'],
        ['Công cụ', 'DASS-42 (sàng lọc) + Phiếu hỏi (biểu hiện & nguyên nhân)'],
        ['Giới tính mẫu', 'Nam 33.5%, Nữ 66.5%'],
    ]
)
doc.add_paragraph()

add_b('Tỷ lệ RLLA:')
add_p('Trên 540 HS được sàng lọc, tỷ lệ mắc RLLA là 14.2% (77 HS). Phân theo mức độ: nhẹ 3.5% (19 HS), vừa 7.2% (39 HS), nặng 2.4% (13 HS), rất nặng 1.1% (6 HS). So với các điều tra dịch tễ trong nước và quốc tế (dao động 8–20% ở thanh thiếu niên), tỷ lệ 14.2% nằm ở mức trung bình. Tỷ lệ RLLA mức độ rất nặng chỉ chiếm 1.1%, cho thấy tình trạng vẫn trong tầm kiểm soát được.')

add_b('Phát hiện bất ngờ — Khối 11 cao nhất:')
add_p('Kết quả sàng lọc cho thấy khối 11 có tỷ lệ RLLA cao nhất (48.1% trên tổng 77 HS), không phải khối 12 như giả thuyết ban đầu. Khối 10 chiếm 31.1%, khối 12 thấp nhất chỉ 20.8%. Ở mức độ vừa và nặng, số HS khối 11 gấp 2–3 lần so với 2 khối còn lại. Lý giải: HS lớp 11 chịu áp lực đa chiều — băn khoăn chọn trường thi, ngành thi, khối thi; quan hệ xã hội phức tạp; thiếu nguồn hỗ trợ thông tin hướng dẫn nên mông lung lo lắng. HS lớp 12 đã xác định được lực học và hướng đi, khả năng tư duy giải quyết vấn đề tốt hơn.')

add_b('Biểu hiện thể chất và tâm lý:')
add_p('85.7% HS có biểu hiện choáng váng, chóng mặt, hoa mắt — đây là biểu hiện phổ biến nhất. 76.6% khó tập trung chú ý, trí nhớ trống rỗng. 62.3% rối loạn giấc ngủ và nóng ruột. Về tâm lý: 85.7% dễ cáu gắt vì chuyện không đáng và dễ yếu mệt mỏi. 81.8% do dự khi phải đưa ra quyết định. 80.5% lo lắng về tình huống bị chế giễu. 67.5% thất vọng về bản thân, cảm thấy vô dụng bất lực. Đặc biệt, mệnh đề "cảm thấy bản thân vô dụng" có tỷ lệ "Rất thường xuyên" cao nhất — nghịch lý ở HS Chuyên vốn giỏi nhưng lại cảm thấy bất lực trước áp lực danh tiếng.')

doc.add_paragraph()

# --- BÀI 3 ---
add_h('Bài 3: Trần Nguyễn Ngọc (2018) — Luận án Tiến sĩ: Hiệu quả điều trị RLLALT bằng liệu pháp thư giãn–luyện tập', 2)

tbl(
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Trần Nguyễn Ngọc — ĐH Y Hà Nội'],
        ['Hướng dẫn', 'PGS.TS Nguyễn Kim Việt'],
        ['Chuyên ngành', 'Tâm thần, Mã số 62720148'],
        ['Địa điểm', 'Viện Sức khỏe Tâm thần, BV Bạch Mai (10/2013–10/2017)'],
        ['Mẫu MT1', '170 bệnh nhân RLLALT (F41.1 theo ICD-10), mô tả cắt ngang'],
        ['Mẫu MT2', '99 bệnh nhân can thiệp thư giãn–luyện tập, 20 buổi/4 tuần, không dùng thuốc'],
        ['Công cụ', 'HAM-A, PSQI, Eysenck EPI, CGI'],
    ]
)
doc.add_paragraph()

add_b('Đặc điểm lâm sàng (n=170):')
add_p('Bệnh nhân phần lớn là nữ (61.8%), tuổi trung bình 43.2 ± 13.6, nhóm 26–45 tuổi thường gặp nhất. Chỉ 22.9% đến đúng chuyên khoa tâm thần lần khám đầu tiên; 26.4% khám tim mạch, 16.5% thần kinh — phản ánh nhận thức yếu về SKTT và kỳ thị bệnh tâm thần tại Việt Nam. Lo âu nặng theo HAM-A chiếm 37.1%. Chủ đề lo âu: gia đình (79.4%), tai nạn/bệnh tật (72.4%), công việc (63.5%). Thường gặp 3 chủ đề lo âu cùng lúc (40%). Triệu chứng phổ biến nhất: khó ngủ vì lo lắng (97%), bồn chồn (93.5%), hồi hộp/tim đập nhanh (89.4%), căng thẳng tâm thần (71.7%).')

add_p('77.2% bệnh nhân có nhân cách hướng nội–không ổn định (Eysenck). 45.3% có sang chấn tâm lý, chủ yếu trường diễn, chủ đề gia đình. 12.4% đồng mắc trầm cảm. Triệu chứng khởi phát thường gặp: hồi hộp/tim đập nhanh (40%), bồn chồn (35.3%), khó ngủ (16.5%). Thời điểm lo âu nặng lên: buổi tối (68.2%). Tần suất xuất hiện trung bình 5.4 lần/tuần, thời gian tồn tại 21–31 phút mỗi lần.')

add_b('Hiệu quả điều trị (n=99):')

tbl(
    ['Chỉ số', 'T0 (trước)', 'T2 (tuần 2)', 'T4 (tuần 4)', 'p (T0–T4)'],
    [
        ['Lo âu nặng (HAM-A)', '45.5%', '22.2%', '11.1%', '<0.0001'],
        ['Tần suất lo âu (lần/tuần)', '5.2 ± 2.7', '3.4 ± 2.6', '1.3 ± 2.0', '<0.0001'],
        ['Số triệu chứng TB', '11.8 ± 3.5', '9.5 ± 3.8', '5.1 ± 4.9', '<0.0001'],
        ['TC thần kinh thực vật', '2.5 ± 1.0', '1.7 ± 1.1', '0.9 ± 1.1', '<0.0001'],
        ['Khó ngủ', '96.9%', '96.9%', '75.7%', '<0.0001'],
        ['Bồn chồn', '96.9%', '96.9%', '52.5%', '<0.0001'],
        ['CGI – thuyên giảm rõ rệt', '1.0%', '—', '41.4%', '<0.0001'],
    ]
)
doc.add_paragraph()

add_p('Liệu pháp thư giãn–luyện tập (kết hợp tự ám thị Schultz + thở khí công + 6 tư thế Yoga) cho thấy hiệu quả rõ rệt trong giảm triệu chứng lo âu và các triệu chứng cơ thể. Tuy nhiên, tác dụng từ từ — không có bệnh nhân nào đạt "cải thiện rõ rệt" hay "cải thiện rất nhiều" theo CGI. Triệu chứng khó ngủ và bồn chồn không thay đổi sau 2 tuần đầu, chỉ cải thiện từ tuần 2–4. Triệu chứng tri giác sai thực tại không cải thiện có ý nghĩa thống kê (p=0.1574). 54.55% bệnh nhân tự đánh giá "có hiệu quả", 12.12% "không hiệu quả".')

doc.add_paragraph()

# ==================== II. ĐIỂM NỔI BẬT ====================
add_h('II. ĐIỂM NỔI BẬT', 1)

pts = [
    ('1. Tỷ lệ RLLA ở HS THPT Việt Nam khá cao (14–18.5%)',
     'Cả 2 bài về học sinh đều xác nhận tỷ lệ RLLA nằm trong dải 14–18.5%, phù hợp với dữ liệu dịch tễ quốc tế (8–20% ở thanh thiếu niên). Đây là bằng chứng quan trọng cho thấy RLLA ở HS THPT Việt Nam là vấn đề sức khỏe tâm thần cần được quan tâm nghiêm túc. Các số liệu này được thu thập trước COVID-19, nên tỷ lệ thực tế hiện nay có thể cao hơn nhiều theo xu hướng toàn cầu.'),
    ('2. Khối 11 có RLLA cao nhất — phát hiện bất ngờ',
     'Bài 2 phát hiện khối 11 có tỷ lệ RLLA cao nhất (48.1%), không phải khối 12 như giả thuyết. Đây là kết quả có giá trị thực tiễn lớn. HS lớp 11 chịu áp lực đa chiều: chưa xác định được trường/ngành thi, quan hệ xã hội phức tạp, thiếu nguồn hỗ trợ. Trong khi lớp 12 đã ổn định hơn về định hướng và tư duy. Kết quả này gợi ý cần tập trung can thiệp tâm lý cho HS lớp 11 từ sớm, không chỉ tập trung vào lớp 12 như thường thấy.'),
    ('3. Nghịch lý HS Chuyên — "bẫy chủ nghĩa hoàn hảo"',
     'HS Chuyên vốn giỏi nhưng lại cảm thấy "vô dụng, bất lực" — tỷ lệ "Rất thường xuyên" cao nhất. Áp lực danh tiếng trường Chuyên trở thành tác nhân gây lo âu thay vì động lực. Đây là hiện tượng "bẫy chủ nghĩa hoàn hảo" (perfectionism trap) được ghi nhận rõ ràng trong bối cảnh giáo dục Việt Nam. Một số HS thuật lại: "Em là HS Chuyên mà mấy cái đó cũng không làm được, em thấy tệ quá." Hiện tượng này cần được nghiên cứu sâu hơn trong bối cảnh áp lực thành tích học tập đặc thù của Việt Nam.'),
    ('4. Yếu tố bản thân > yếu tố học tập trong gây RLLA',
     'Bài 1 chỉ ra nhóm yếu tố "bản thân học sinh" chi phối RLLA mạnh nhất (r=0.42, 85.4%), vượt xa yếu tố học tập (r=0.37, 74%). Thiếu kỹ năng ứng phó, thiếu giao tiếp chia sẻ là yếu tố nội tại quan trọng nhất. Điều này gợi ý rằng can thiệp kỹ năng sống (khả năng giải quyết vấn đề, giao tiếp, điều hòa cảm xúc) có thể hiệu quả hơn chỉ giảm tải học tập. Đây là phát hiện có ý nghĩa thực tiễn cho việc xây dựng chương trình tâm lý học đường.'),
    ('5. Luận án tiến sĩ — công trình bài bản nhất',
     'Bài 3 là nghiên cứu đầu tiên tại Việt Nam đánh giá hiệu quả liệu pháp thư giãn–luyện tập trên RLLALT theo ICD-10, với cỡ mẫu 170 bệnh nhân. Phần tổng quan bệnh sinh rất xuất sắc, trình bày chi tiết cơ chế GABA, Serotonin, Norepinephrine, Neurosteroid, mạch thần kinh với nhiều tham khảo quốc tế. Đây là đóng góp lý thuyết quan trọng cho nghiên cứu tâm thần học tại Việt Nam. Phương pháp can thiệp 3 phần (thư giãn + khí công + Yoga) là đặc trưng riêng của trường phái tâm thần học Việt Nam.'),
    ('6. Liệu pháp hiệu quả nhưng tác dụng chậm',
     'Liệu pháp thư giãn–luyện tập giảm lo âu nặng từ 45.5% xuống 11.1% sau 4 tuần. Tuy nhiên, không có bệnh nhân nào đạt "cải thiện rõ rệt" theo CGI — liệu pháp tác dụng từ từ. Phù hợp với bản chất RLLALT là bệnh mạn tính. Bồn chồn và khó ngủ chỉ cải thiện từ tuần 2–4, không thay đổi 2 tuần đầu. Gợi ý cần kéo dài liệu trình >4 tuần và kết hợp với các phương pháp khác (CBT, thuốc) cho trường hợp nặng.'),
    ('7. Chỉ 22.9% đến đúng chuyên khoa tâm thần',
     'Hầu hết bệnh nhân khám tim mạch (26.4%) hoặc thần kinh (16.5%) trước khi đến tâm thần. Phản ánh 3 vấn đề: (1) nhận thức yếu về SKTT ở cả bệnh nhân và nhân viên y tế, (2) kỳ thị bệnh tâm thần khiến người bệnh tránh né, (3) triệu chứng cơ thể của RLLALT dễ gây nhầm lẫn với bệnh lý thực thể. Cần đào tạo thêm cho bác sĩ đa khoa về nhận biết RLLALT và giảm kỳ thị thông qua truyền thông sức khỏe cộng đồng.'),
]

for t, b in pts:
    add_b(t)
    add_p(b)

doc.add_paragraph()

# ==================== III. PHẢN BIỆN ====================
add_h('III. PHẢN BIỆN', 1)

add_h('Bài 1 (Nguyễn Thị Vân):', 2)
add_p('Mẫu thuận tiện (n=90 cho phần phân tích yếu tố) không đại diện, có thể có thiên lệch chọn mẫu. Không thể tổng quát hóa cho các tỉnh khác ngoài TP.HCM. Thang đo STAI thích nghi từ năm 2000 — đã hơn 20 năm, cần cập nhật đặc tính đo lường (psychometric properties). Thiếu phân tích đa biến (multivariate) để kiểm soát yếu tố gây nhiễu giữa các nhóm. Con số 85.4% giải thích biến thiên từ yếu tố bản thân là quá cao — cần kiểm tra lại phương pháp hồi quy vì có thể overfitting hoặc multicollinearity. Tài liệu tham khảo chỉ có 4 nguồn, quá ít cho một bài nghiên cứu. Không có nhóm đối chứng (HS không có RLLA) để so sánh.')

add_h('Bài 2 (Trần Thị Mỵ Lương):', 2)
add_p('Chỉ nghiên cứu tại 1 trường Chuyên — đặc thù cao, không đại diện cho HS THPT nói chung. Tỷ lệ nữ 66.5% trong mẫu gây thiên lệch giới, có thể ảnh hưởng kết quả vì nữ thường có tỷ lệ lo âu cao hơn nam. Sử dụng DASS-42 để sàng lọc nhưng không nêu rõ điểm cắt (cutoff score) được sử dụng. Phần phỏng vấn sâu và phiếu hỏi trên 77 HS: không nêu rõ phương pháp phân tích thống kê cụ thể. Thiếu nhóm đối chứng (HS không có RLLA) để so sánh biểu hiện. Kết luận về khối 11 cần được kiểm chứng trên mẫu lớn hơn, đa trường. Phần lý giải nguyên nhân chưa được hỗ trợ bằng dữ liệu định lượng mà chủ yếu là mô tả định tính.')

add_h('Bài 3 (Trần Nguyễn Ngọc — Luận án TS):', 2)
add_p('Hạn chế lớn nhất: KHÔNG CÓ NHÓM CHỨNG. Thiết kế trước–sau không có nhóm đối chứng không thể loại trừ hiệu ứng placebo, hồi quy về trung bình, hoặc tác động của môi trường nằm viện (nội trú). 41.8% bệnh nhân bỏ cuộc (170 → 99 hoàn thành) — tỷ lệ bỏ cuộc rất cao, có thể gây thiên lệch sống sót (survival bias); những người bỏ cuộc có thể là nhóm nặng nhất hoặc ít tuân thủ nhất. Không mù (no blinding): bác sĩ đánh giá CGI cũng là người hướng dẫn điều trị → thiên lệch đánh giá. Chỉ theo dõi 4 tuần — quá ngắn cho RLLALT (bệnh mạn tính, 60% bệnh nhân hồi phục trong 12 năm theo Bruce et al.). Liệu pháp chỉ thực hiện tại 1 cơ sở (Viện SKTT, BV Bạch Mai) nên không tổng quát hóa được. Tuy nhiên, phần tổng quan bệnh sinh RLLALT rất xuất sắc và đây là đóng góp lý thuyết quan trọng.')

doc.add_paragraph()

# ==================== IV. HƯỚNG NGHIÊN CỨU ====================
add_h('IV. HƯỚNG NGHIÊN CỨU ĐỂ HOÀN THIỆN VÀ PHÁT TRIỂN', 1)

add_h('A. Mở rộng nghiên cứu dịch tễ:', 2)
add_p('1. Nghiên cứu cắt ngang đa trung tâm về RLLA ở HS THPT tại Việt Nam — bao gồm nông thôn, thành thị, miền núi. Mẫu cần >3000 HS, sử dụng GAD-7 hoặc DASS-21 (đã chuẩn hóa quốc tế). Việc sử dụng công cụ chuẩn hóa giúp so sánh trực tiếp với dữ liệu quốc tế và tăng độ tin cậy của kết quả.')
add_p('2. So sánh tỷ lệ RLLA trước và sau COVID-19 ở HS Việt Nam. Hiện chưa có dữ liệu nào về vấn đề này. Các nghiên cứu quốc tế cho thấy lo âu ở trẻ em/vị thành niên tăng gấp đôi trong đại dịch. Cần có dữ liệu tương tự cho Việt Nam để xây dựng chính sách hỗ trợ sức khỏe tâm thần học đường.')
add_p('3. Nghiên cứu dọc theo dõi HS từ lớp 10 đến lớp 12 để xác nhận phát hiện "khối 11 cao nhất". Thiết kế này cho phép theo dõi sự thay đổi lo âu theo thời gian và xác định thời điểm can thiệp tối ưu. Cần ít nhất 2–3 trường (thường + chuyên) để so sánh sự khác biệt giữa các loại hình trường.')

add_h('B. Yếu tố nguy cơ đặc thù Việt Nam:', 2)
add_p('4. Vai trò của áp lực thi cử, học thêm, kỳ vọng gia đình — cần nghiên cứu RCT so sánh nhóm có can thiệp giảm áp lực vs. nhóm đối chứng. Việt Nam có văn hóa học tập cạnh tranh cao, học thêm phổ biến. Cần lượng hóa mối liên hệ giữa số giờ học thêm và mức độ lo âu. Đây là đề tài có tính ứng dụng thực tiễn cao cho chính sách giáo dục.')
add_p('5. Tác động của mạng xã hội lên lo âu ở HS Việt Nam — đề tài mới, chưa có nghiên cứu. Trên thế giới, sử dụng >7 nền tảng MXH tăng nguy cơ lo âu gấp 3 lần. Cần khảo sát thời gian sử dụng điện thoại/MXH và tương quan với triệu chứng lo âu trên mẫu HS Việt Nam. Tổng Y sĩ Hoa Kỳ đã ra khuyến cáo chính thức năm 2023 về vấn đề này.')
add_p('6. Yếu tố văn hóa: vai trò của "thể diện" (face), chủ nghĩa tập thể, kỳ thị bệnh tâm thần — so sánh với dữ liệu quốc tế. Việt Nam có bối cảnh văn hóa đặc thù ảnh hưởng đến cách biểu hiện và tìm kiếm hỗ trợ cho lo âu. Cần nghiên cứu định tính kết hợp định lượng để hiểu sâu các cơ chế văn hóa–xã hội.')

add_h('C. Can thiệp và điều trị:', 2)
add_p('7. RCT so sánh liệu pháp thư giãn–luyện tập với CBT hoặc với thuốc (SSRI). Thiếu nhóm chứng là hạn chế lớn nhất của bài 3 cần được khắc phục. Thiết kế 3 nhánh (thư giãn vs CBT vs thuốc) với follow-up 6 tháng sẽ cung cấp bằng chứng mạnh nhất cho thực hành lâm sàng tại Việt Nam.')
add_p('8. Kéo dài theo dõi liệu pháp thư giãn–luyện tập đến 3–6 tháng. Đánh giá tỷ lệ tái phát. RLLALT là bệnh mạn tính (60% bệnh nhân hồi phục trong 12 năm theo Bruce et al.), nên theo dõi ngắn hạn 4 tuần chưa đủ để kết luận về hiệu quả lâu dài và khả năng duy trì kết quả điều trị.')
add_p('9. Phát triển chương trình CBT phiên bản Việt Nam cho HS THPT. Hiện tại CBT chưa phổ biến ở Việt Nam dù đã được chứng minh là liệu pháp hàng đầu trên thế giới cho rối loạn lo âu. Cần dịch và thích nghi hóa protocol CBT, đào tạo nhân lực, và thử nghiệm tại trường học trong bối cảnh Việt Nam.')
add_p('10. Ứng dụng kỹ thuật số: App thư giãn/chánh niệm bằng tiếng Việt cho HS. Thị trường DTx toàn cầu đạt 3.23 tỷ USD (2021). Cần phát triển ứng dụng phiên bản Việt với nội dung dựa trên bằng chứng, có đánh giá hiệu quả bằng RCT. Đây là hướng tiếp cận chi phí thấp, có thể nhân rộng nhanh trong hệ thống trường học.')

add_h('D. Phương pháp luận:', 2)
add_p('11. Sử dụng thiết kế RCT hoặc quasi-experimental có nhóm chứng. Đây là yêu cầu tối thiểu để nâng cấp mức độ bằng chứng. Cần tính cỡ mẫu đủ lớn, phân bố ngẫu nhiên, và mù đánh giá (assessor blinding) để giảm thiên lệch trong nghiên cứu can thiệp.')
add_p('12. Sử dụng công cụ đo lường chuẩn hóa quốc tế (GAD-7, PHQ-9, DASS-21) thay vì chỉ STAI. GAD-7 có sensitivity 0.81, specificity 0.78 qua meta-analysis 43 nghiên cứu. Việc sử dụng công cụ chuẩn hóa giúp so sánh trực tiếp với văn bản quốc tế và tăng tính giá trị ngoại của nghiên cứu.')
add_p('13. Cần phân tích đa biến (logistic regression, SEM) để xác định yếu tố nguy cơ độc lập. Các bài hiện tại chủ yếu dùng thống kê mô tả và tương quan đơn biến, chưa kiểm soát yếu tố gây nhiễu. SEM sẽ cho phép kiểm định mô hình nguyên nhân–kết quả phức tạp hơn, phù hợp với bản chất đa yếu tố của RLLA.')

doc.add_paragraph()

# ==================== V. DANH SÁCH BÀI CẦN ĐỌC ====================
add_h('V. DANH SÁCH BÀI CẦN ĐỌC LIÊN QUAN', 1)

# Helper: table with custom column widths (in cm)
def tbl_w(headers, rows, widths_cm):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Set column widths
    for i, w in enumerate(widths_cm):
        for row in t.rows:
            row.cells[i].width = Cm(w)
    # Header
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    # Data
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
    return t

# Widths: # (0.6), Tác giả (2.5), Năm (0.8), Tạp chí (2.8), Nội dung (9.3) = ~16cm for A4
W = [0.6, 2.5, 0.8, 2.8, 9.3]

add_h('A. Tổng quan & lý thuyết nền tảng:', 2)
tbl_w(
    ['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['1', 'Penninx BWJH, Pine DS, Holmes EA, Reif A', '2021', 'The Lancet (IF~169)',
         'Bài seminar hàng đầu về rối loạn lo âu. Trình bày tổng quan dịch tễ, cơ chế bệnh sinh (mạch thần kinh, di truyền, biểu sinh), chẩn đoán, và điều trị. Nhấn mạnh CBT và hợp chất serotonergic là lựa chọn hàng đầu. Chỉ ra rằng chưa có biện pháp phòng ngừa bền vững; kháng trị vẫn là thách thức đòi hỏi y học chính xác. ~800+ trích dẫn.'],
        ['2', 'Szuhany KL, Simon NM', '2022', 'JAMA (IF~157)',
         'Tổng quan toàn diện trên JAMA về GAD, SAD, và rối loạn hoảng sợ. Rối loạn lo âu ảnh hưởng ~34% người Mỹ trong suốt cuộc đời. Điều trị hàng đầu: SSRI (sertraline), SNRI (venlafaxine ER), và CBT. Bài viết đề cập chẩn đoán, dịch tễ, sinh bệnh học tại cơ sở chăm sóc ban đầu. ~259+ trích dẫn.'],
        ['3', 'Craske MG, Stein MB, Eley TC, Milad MR, Holmes A, Rapee RM, Wittchen HU', '2017', 'Nature Rev Disease Primers (IF~65)',
         'Bài primer nền tảng của Nature Reviews, bao phủ toàn bộ lĩnh vực: dịch tễ, chẩn đoán, sinh bệnh học (mạch thần kinh, di truyền, thiên lệch nhận thức), yếu tố nguy cơ, phòng ngừa, điều trị. Công trình này đã định hướng nghiên cứu lo âu cho đến hiện tại. ~1500+ trích dẫn — một trong những bài được trích dẫn nhiều nhất trong lĩnh vực.'],
    ], W
)
doc.add_paragraph()

add_h('B. Lo âu ở học sinh (liên quan bài 1 & 2):', 2)
tbl_w(
    ['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['4', 'Racine N, McArthur BA, Cooke JE, et al.', '2021', 'JAMA Pediatrics (IF~26)',
         'Meta-analysis toàn cầu về trầm cảm và lo âu ở trẻ em/vị thành niên trong COVID-19. Tỷ lệ lo âu tăng cao: 20.5%, trầm cảm 25.2% — gấp đôi so với trước đại dịch. Tỷ lệ tăng theo thời gian và cao hơn ở nữ, vị thành niên lớn tuổi hơn. ~3000+ trích dẫn — bài được trích dẫn nhiều nhất về chủ đề này.'],
        ['5', 'Samji H, Wu J, et al.', '2023', 'JAMA Pediatrics (IF~26)',
         'Meta-analysis 53 nghiên cứu dọc, 12 quốc gia (n=40,807) so sánh lo âu/trầm cảm trước vs trong COVID. Bằng chứng mạnh cho thấy trầm cảm tăng trong đại dịch, đặc biệt ở nữ và các nước thu nhập cao. Lo âu tăng nhẹ nhưng có ý nghĩa. Hướng dẫn cho chính sách can thiệp hậu COVID.'],
        ['6', 'Xu Q, Mao Z, et al.', '2021', 'J Affective Disorders (IF~6.6)',
         'Khảo sát cắt ngang quy mô cực lớn trên 373,216 HS THCS và THPT trong dịch COVID-19 tại Trung Quốc. Tỷ lệ lo âu chung: 9.89%. THCS > THPT (13.89% vs 12.93%). Nông thôn > thành thị (11.33% vs 8.77%). Tuổi, giới, nơi cư trú, mức lo lắng và sợ hãi là yếu tố liên quan. Một trong những khảo sát lo âu học sinh lớn nhất toàn cầu.'],
        ['7', 'Chen Z, et al.', '2023', 'BMC Psychiatry (IF~4.4)',
         'Nghiên cứu cắt ngang trên HS THCS tại Zigong (Tây Trung Quốc) sử dụng PHQ-9, GAD-7, thang đo nạn nhân bắt nạt đa chiều, chỉ số chất lượng giấc ngủ Pittsburgh, và thang đo rối loạn chơi game. Trầm cảm/lo âu ảnh hưởng ~1/4 HS. Kiểm tra yếu tố nhân khẩu, gia đình, trường học, lối sống, hành vi. Phù hợp so sánh với bối cảnh VN.'],
        ['8', 'Qiu Z, Guo Y, Wang J, Zhang H', '2022', 'Frontiers in Psychology (IF~3.8)',
         'Nghiên cứu cắt ngang về mối liên hệ giữa phong cách nuôi dạy, sức bền (resilience) với trầm cảm và lo âu ở HS THCS Trung Quốc. Mô hình "nuôi dạy tiêu cực" liên quan tích cực với lo âu; "nuôi dạy tích cực" liên quan nghịch. Sức bền đóng vai trò yếu tố bảo vệ trung gian. Phù hợp so sánh trực tiếp với bài 1 (Nguyễn Thị Vân) về yếu tố gia đình.'],
    ], W
)
doc.add_paragraph()

add_h('C. Liệu pháp thư giãn & mindfulness (liên quan bài 3):', 2)
tbl_w(
    ['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['9', 'Papola D, Miguel C, et al.', '2024', 'JAMA Psychiatry (IF~25)',
         'Tổng quan hệ thống và phân tích meta mạng lưới 65 RCT với >5,000 người lớn mắc GAD. CBT thế hệ thứ ba, CBT tiêu chuẩn, và liệu pháp thư giãn đều hiệu quả hơn chăm sóc thông thường. CBT nên là lựa chọn hàng đầu. Tỷ lệ bỏ trị tương đương giữa các phương pháp. Cung cấp bằng chứng mạnh nhất hiện tại cho lựa chọn liệu pháp tâm lý điều trị GAD.'],
        ['10', '(Nhiều tác giả)', '2023', 'Nature Mental Health',
         'Tổng quan hệ thống và meta-analysis dữ liệu cá nhân (IPD) các RCT về chương trình dựa trên chánh niệm (MBI). Chứng minh kích thước hiệu ứng trung bình–mạnh cho giảm lo âu và trầm cảm. Hiệu quả MBI tương đương với CBT trong giảm triệu chứng lo âu. Hỗ trợ cơ sở lý thuyết cho phần luyện thư giãn trong bài 3.'],
        ['11', 'Manzoni GM, et al.', '2008', 'BMC Psychiatry',
         'Tổng quan hệ thống và meta-analysis 10 năm (1997–2007) trên 27 nghiên cứu về huấn luyện thư giãn cho lo âu. Xác nhận các liệu pháp thư giãn dựa trên tự ám thị có thể điều trị triệu chứng lo âu. Đây là bài tham khảo trực tiếp cho cơ sở lý thuyết của liệu pháp thư giãn–luyện tập trong bài 3 (Trần Nguyễn Ngọc).'],
        ['12', 'Singh B, Olds T, Curtis R, et al.', '2023', 'Br J Sports Med (IF~18)',
         'Umbrella review tổng hợp 97 tổng quan hệ thống (1,039 thử nghiệm, >128,000 người). Vận động thể chất có hiệu ứng trung bình đối với lo âu (median effect size = -0.42). Cường độ cao hơn liên quan đến cải thiện lớn hơn. Tập thể dục "nên là phương pháp chủ đạo trong quản lý trầm cảm, lo âu và đau khổ tâm lý". ~500+ trích dẫn.'],
    ], W
)
doc.add_paragraph()

add_h('D. Công cụ sàng lọc & bối cảnh Việt Nam:', 2)
tbl_w(
    ['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['13', '(Nhiều tác giả)', '2025', 'J Affective Disorders (IF~6.6)',
         'Meta-analysis cập nhật 43 nghiên cứu xác nhận giá trị dự đoán của GAD-7 và GAD-2 cho sàng lọc rối loạn lo âu. GAD-7: độ nhạy 0.81, độ đặc hiệu 0.78, AUC 0.87. Tại ngưỡng 8: độ nhạy 0.83, độ đặc hiệu 0.84. Khuyến nghị sử dụng GAD-7 thay STAI trong nghiên cứu tại Việt Nam để tăng tính so sánh quốc tế.'],
        ['14', 'Walter HJ, Bukstein OG, et al. (AACAP)', '2020', 'JAACAP (IF~12)',
         'Hướng dẫn thực hành lâm sàng chính thức của Viện Hàn lâm AACAP về đánh giá và điều trị lo âu ở trẻ em/vị thành niên. CBT và SSRI có bằng chứng mạnh nhất. Thử nghiệm CAMS (488 trẻ): kết hợp CBT + sertraline hiệu quả nhất (81%) vs CBT đơn lẻ (60%) vs sertraline đơn lẻ (55%). Tham khảo quan trọng cho xây dựng protocol điều trị tại VN.'],
        ['15', 'Hồ Hữu Tính, Nguyễn Doãn Thành', '2010', 'Tạp chí Y học TP HCM',
         'Nghiên cứu thực trạng stress, lo âu và các yếu tố liên quan ở HS cấp 3 trường THPT Phan Bội Châu, Phan Thiết, Bình Thuận. Đây là một trong số ít nghiên cứu về lo âu HS THPT tại Việt Nam được công bố trên tạp chí chuyên ngành. Được trích dẫn trong bài 2 (Trần Thị Mỵ Lương) làm cơ sở so sánh.'],
        ['16', 'Nguyễn Phước Bình', '2010', 'Luận văn CKII, ĐH Y Hà Nội',
         'Nghiên cứu đặc điểm lâm sàng RLLALT trên 90 bệnh nhân nội trú tại Viện SKTT, BV Bạch Mai. Được trích dẫn làm cơ sở so sánh chính trong bài 3 (Trần Nguyễn Ngọc). Kết quả: RLLALT thường gặp ở nhóm lao động chân tay (68.9%), triệu chứng hồi hộp/tim đập nhanh phổ biến nhất, sang chấn tâm lý chủ đề xã hội chiếm 92%.'],
    ], W
)

doc.add_paragraph()
add_i('Báo cáo này được viết ngày 25/03/2026. Dựa trên phân tích 3 bài PDF trong thư mục papers/Viet-nam/.')

output = r'c:\Users\HLC\OneDrive\read_books\Lo-au\BAO_CAO_3_BAI_VIET_NAM.docx'
doc.save(output)
print(f'OK: {output}')
