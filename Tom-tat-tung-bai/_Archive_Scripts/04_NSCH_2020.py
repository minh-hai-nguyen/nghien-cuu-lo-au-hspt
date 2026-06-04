# -*- coding: utf-8 -*-
"""Bài 4: NSCH 2020 — theo phong cách CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 4 =====
bl('Tóm tắt bài 4', bold=True)

# BƯỚC 1: ĐỊNH DANH
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình \u00ab Khảo sát Quốc gia về Sức khỏe Trẻ em \u00bb (National Survey of Children\u2019s Health \u2014 NSCH), do Data Resource Center for Child & Adolescent Health thực hiện, được tài trợ bởi Maternal and Child Health Bureau (HRSA) phối hợp với Cục Điều tra Dân số Hoa Kỳ (U.S. Census Bureau), khảo sát trên mẫu 55.162 trẻ em và thanh thiếu niên từ 0 đến 17 tuổi (trong đó 18.397 thanh thiếu niên 12\u201317 tuổi) trên toàn bộ 50 tiểu bang và Washington D.C., Hoa Kỳ, dữ liệu năm 2023 (khảo sát hàng năm từ 2016\u20132024).')

# BƯỚC 2: TỔNG QUAN PP + KT1
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng bảng câu hỏi khảo sát qua đường bưu điện và trực tuyến, do cha mẹ/người chăm sóc hoàn thành thay mặt cho trẻ, dựa trên tình trạng đã được chẩn đoán bởi bác sĩ/chuyên gia y tế ("current, diagnosed condition"). Nói cách khác, đây là phương pháp khảo sát quốc gia dựa trên chẩn đoán lâm sàng, không phải sàng lọc triệu chứng.', bold=True)

# BƯỚC 3: ĐỊNH NGHĨA (KT5)
bl('NSCH là khảo sát thường niên quy mô quốc gia do Cục Điều tra Dân số Hoa Kỳ quản lý, cung cấp dữ liệu về sức khỏe thể chất và tâm thần, khả năng tiếp cận chăm sóc y tế, bối cảnh gia đình, trường học và cộng đồng của trẻ em. Dữ liệu được cân trọng số (weighted) để đại diện cho dân số trẻ em Hoa Kỳ từ 0\u201317 tuổi không sống trong cơ sở tập trung.')

# BƯỚC 4: BIỆN MINH (KT4)
bl('Tổng quan tài liệu của chúng tôi cho thấy đa số các nghiên cứu về sức khỏe tâm thần thanh thiếu niên sử dụng thang sàng lọc triệu chứng (PHQ-9, GAD-7, DASS-21), cho kết quả tỷ lệ cao hơn nhiều so với phương pháp chẩn đoán lâm sàng. NSCH là một trong số ít nguồn dữ liệu quốc gia sử dụng phương pháp dựa trên chẩn đoán, cho phép ước tính tỷ lệ rối loạn thực sự chứ không chỉ triệu chứng.', bold=True)

# BƯỚC 5: LIỆT KÊ DỮ LIỆU
bl('Dữ liệu bao gồm sức khỏe thể chất và cảm xúc, chất lượng chăm sóc y tế tại nhà (medical home), tương tác gia đình, sức khỏe cha mẹ, trải nghiệm trường học, an toàn khu phố, và các tình trạng sức khỏe tâm thần được chẩn đoán hiện tại (lo âu, trầm cảm, vấn đề hành vi/ứng xử).')

# BƯỚC 6: MÔ TẢ QUY TRÌNH
bl('Khảo sát được thực hiện hàng năm từ 2016 đến 2024 dưới dạng bảng câu hỏi gửi qua đường bưu điện và trực tuyến. Cha mẹ hoặc người chăm sóc chính hoàn thành bảng câu hỏi thay mặt cho một trẻ được chọn ngẫu nhiên trong hộ gia đình. Dữ liệu được cân trọng số theo mức nghèo gia đình, loại bảo hiểm y tế và các đặc điểm nhân khẩu học để đảm bảo tính đại diện quốc gia.')

# BƯỚC 9: ĐỊNH LƯỢNG — BẢNG
rb('Kết quả nghiên cứu định lượng')

bl('Tỷ lệ các tình trạng sức khỏe tâm thần được chẩn đoán ở thanh thiếu niên 12\u201317 tuổi (dữ liệu 2023, N = 18.397):', bold=True)
tbl(
    ['Tình trạng', 'Tỷ lệ chung', 'Nữ', 'Nam'],
    [
        ['Bất kỳ tình trạng SKTT nào', '20,3%', '\u2014', '\u2014'],
        ['Lo âu (phổ biến nhất)', '16,1%', '20,1%', '12,3%'],
        ['Trầm cảm', '8,4%', '10,9%', '6,0%'],
        ['Vấn đề hành vi/ứng xử', '6,3%', '4,3%', '8,2%'],
    ],
    widths=[4.5, 2.5, 2.5, 2.5]
)
doc.add_paragraph()

bl('Xu hướng 7 năm (2016\u20132023):', bold=True)
tbl(
    ['Chỉ số', '2016', '2023', 'Thay đổi'],
    [
        ['Tổng tình trạng SKTT', '15,0%', '20,3%', '+35%'],
        ['Lo âu', '10,0%', '16,1%', '+61%'],
        ['Trầm cảm', '5,8%', '8,4%', '+45%'],
        ['Vấn đề hành vi/ứng xử', '\u2014', '6,3%', 'Ổn định'],
    ],
    widths=[4.5, 2.5, 2.5, 2.5]
)
doc.add_paragraph()

bl('Tác động đến trường học và xã hội:', bold=True)
tbl(
    ['Tác động', 'Nguy cơ so với nhóm không có SKTT'],
    [
        ['Khó khăn kết bạn/duy trì tình bạn', 'Cao gấp 10 lần'],
        ['Nghỉ học ≥11 ngày', 'Cao gấp 5 lần'],
        ['Trường liên hệ về vấn đề', 'Cao gấp 4 lần'],
        ['Không tham gia học tập', 'Cao gấp 3 lần'],
        ['Bị bắt nạt', 'Cao gấp 2 lần'],
    ],
    widths=[6.0, 6.0]
)
doc.add_paragraph()

# BƯỚC 10: ĐỐI CHIẾU (KT3)
bh3('Khác biệt giới tính : Nữ giới luôn được báo cáo có tỷ lệ lo âu và trầm cảm cao hơn nam giới trong y văn quốc tế (Salk và cộng sự, 2017). Thực tế, dữ liệu NSCH xác nhận xu hướng này ở cấp quốc gia Hoa Kỳ: nữ 20,1% so với nam 12,3% về lo âu, nữ 10,9% so với nam 6,0% về trầm cảm. Tuy nhiên, nam giới có tỷ lệ vấn đề hành vi/ứng xử cao hơn (8,2% so với 4,3%) \u2014 gợi ý rằng sức khỏe tâm thần biểu hiện khác nhau theo giới: nữ nội hóa (lo âu, trầm cảm), nam ngoại hóa (hành vi).')

# BƯỚC 11+12: NHẬN XÉT
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về xu hướng gia tăng và khoảng cách chẩn đoán-sàng lọc')
bl('*Xu hướng gia tăng lo âu.* Nổi bật nhất là mức tăng 61% của lo âu được chẩn đoán trong 7 năm (2016\u20132023), từ 10,0% lên 16,1%. Đây là tỷ lệ tăng nhanh nhất trong các tình trạng sức khỏe tâm thần thanh thiếu niên tại Hoa Kỳ. Sự gia tăng này có thể phản ánh cả tăng thực sự trong tỷ lệ mắc bệnh lẫn tăng nhận thức và giảm kỳ thị (destigmatization), dẫn đến nhiều gia đình tìm kiếm chẩn đoán hơn.')
bl('*Khoảng cách chẩn đoán và sàng lọc.* So sánh tỷ lệ lo âu 16,1% (NSCH, dựa trên chẩn đoán) với 40,6% (Hoa, 2024, GAD-7, Việt Nam), 54,4% (Ngo Anh Vinh, 2024, DASS-21, Việt Nam) và chỉ 2,3% (V-NAMHS, 2022, DISC-5/DSM-5, Việt Nam) cho thấy phương pháp đo lường quyết định mạnh mẽ đến kết quả: sàng lọc cho tỷ lệ cao gấp 2,5\u201325 lần so với chẩn đoán.')
bl('*Tác động đến giáo dục.* Phát hiện đáng chú ý là thanh thiếu niên có tình trạng sức khỏe tâm thần gặp khó khăn kết bạn cao gấp 10 lần và nghỉ học ≥11 ngày cao gấp 5 lần. Điều này tương tự kết luận của Jenkins và cộng sự (2023) về ảnh hưởng của sức khỏe tâm thần đến sự tham gia giáo dục.')
bl('*Tiếp cận điều trị.* 82,6% thanh thiếu niên cần dịch vụ đã nhận điều trị, nhưng 61,0% gặp khó khăn trong tiếp cận (tăng 35% từ 2018). Đây là một tỷ lệ tiếp cận cao hơn đáng kể so với các nước đang phát triển: Việt Nam chỉ 8,4% (V-NAMHS, 2022) và Philippines khoảng 2% (Puyat, 2025).')

# BƯỚC 13: KẾT LUẬN (Kiến trúc B)
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy lo âu là tình trạng sức khỏe tâm thần phổ biến nhất (16,1%) ở thanh thiếu niên Hoa Kỳ với mức tăng 61% trong 7 năm \u2014 dữ liệu từ khảo sát quốc gia quy mô lớn nhất (n = 55.162), gợi ý rằng lo âu đang trở thành cuộc khủng hoảng sức khỏe cộng đồng ở thanh thiếu niên. Sự khác biệt giới tính (nữ 20,1% so với nam 12,3%) và tác động nghiêm trọng đến giáo dục (khó kết bạn gấp 10 lần, nghỉ học gấp 5 lần) nhấn mạnh nhu cầu can thiệp sớm. Trong bối cảnh quốc gia phát triển với hệ thống chẩn đoán tốt như Hoa Kỳ mà 61% vẫn gặp khó khăn tiếp cận, thách thức ở các nước đang phát triển như Việt Nam (8,4% tiếp cận) là rất lớn.', bold=True)

# PHẢN BIỆN
rh2('Phản biện')
bl('Điểm mạnh nổi bật là cỡ mẫu rất lớn (>50.000 hàng năm) và tính đại diện quốc gia nhờ hệ thống trọng số thống kê. Tuy nhiên, dữ liệu được thu thập từ cha mẹ/người chăm sóc chứ không phải từ thanh thiếu niên trực tiếp \u2014 nghiên cứu cho thấy cha mẹ thường đánh giá thấp triệu chứng nội hóa (lo âu, trầm cảm). Do đó, tỷ lệ 16,1% có thể là ước tính thiếu (underestimate). Phương pháp dựa trên chẩn đoán phụ thuộc vào khả năng tiếp cận dịch vụ \u2014 trẻ từ gia đình thu nhập thấp, nông thôn hoặc thiểu số có thể chưa bao giờ được chẩn đoán dù có triệu chứng. Xu hướng tăng 61% lo âu (2016\u20132023) cần diễn giải thận trọng: có thể phản ánh tăng nhận thức và giảm kỳ thị song song với tăng thực sự. Phương pháp bưu điện/trực tuyến có thiên lệch tự chọn (self-selection bias).')

# HƯỚNG NC
rh2('Hướng nghiên cứu tiếp theo')
bl('Bổ sung tự báo cáo trực tiếp từ thanh thiếu niên song song với báo cáo của cha mẹ để so sánh và giảm ước tính thiếu. Phân tích sâu nhóm chưa được chẩn đoán nhưng có triệu chứng \u2014 kết hợp sàng lọc PHQ-9/GAD-7 với dữ liệu chẩn đoán NSCH trong cùng mẫu. Nghiên cứu nguyên nhân xu hướng tăng 61% lo âu: phân tách tỷ lệ "tăng phát hiện" và "tăng thực sự". Mở rộng phân tích tác động COVID-19 trên dữ liệu dọc 2019\u20132024.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50 Cao. Cỡ mẫu quốc gia >50.000, phương pháp dựa trên chẩn đoán lâm sàng, theo dõi xu hướng 7 năm, nhưng hạn chế bởi báo cáo gián tiếp qua cha mẹ.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('04_NSCH_2020.docx')
print('04_NSCH_2020.docx OK')
