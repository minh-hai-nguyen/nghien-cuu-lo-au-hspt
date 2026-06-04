# -*- coding: utf-8 -*-
"""Sinh file .docx bao cao ra soat 21 tai lieu tham khao + luan diem than bai
cua ban thao bai2 KHGDVN (build_bai2_KHGDVN.py). 18/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Bao_cao_ra_soat_TLTK_bai2_KHGDVN_18052026.docx')
PAGE_W = 15.5
RED = RGBColor(0xCC, 0, 0)

GREEN_SH = 'C6E0B4'
YELLOW_SH = 'FFE699'
RED_SH = 'F4B0B0'

# ---------- helpers ----------
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)

def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(doc, text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def bullet(doc, text, bold_prefix=None, size=12, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = True
        if color is not None: r.font.color.rgb = color
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

def add_hyperlink(paragraph, url, text, size=11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def ref_item(doc, text, doi_label=None, doi_url=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    if doi_label and doi_url:
        sep = p.add_run('  '); sep.font.name = 'Times New Roman'; sep.font.size = Pt(11)
        add_hyperlink(p, doi_url, doi_label, size=11)

def cell_fmt(cell, size, bold=False, color=None, align='justify'):
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY if align == 'justify' else WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold
            if color is not None: r.font.color.rgb = color

def status_shade(status):
    s = status.upper()
    if s.startswith('ĐÚNG'):
        return GREEN_SH
    if s.startswith('TẠM'):
        return YELLOW_SH
    return RED_SH

def audit_table(doc, headers, rows, widths, hdr_size=9, body_size=8.5):
    assert abs(sum(widths) - PAGE_W) < 0.1, sum(widths)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        cell_fmt(c, hdr_size, bold=True, align='center')
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            cell_fmt(c, body_size)
    return t

# ============================================================
doc = make_doc()

# ---- Tieu de ----
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('BÁO CÁO RÀ SOÁT TÀI LIỆU THAM KHẢO VÀ LUẬN ĐIỂM\n'
                'BẢN THẢO BÀI BÁO bai2 — TẠP CHÍ KHOA HỌC GIÁO DỤC VIỆT NAM')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Đối tượng rà soát: file 06_Scripts/build_bai2_KHGDVN.py — '
                'Ngày lập: 18/05/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

# ---- 1. Boi canh ----
H(doc, '1. Bối cảnh và phạm vi rà soát', 1)
P(doc, 'Bản thảo bai2 ("Khoảng trống can thiệp rối loạn lo âu ở học sinh trung học cơ sở") '
       'là một bài tổng quan tài liệu chuẩn bị nộp Tạp chí Khoa học Giáo dục Việt Nam. Trong '
       'quá trình sửa một lỗi trích dẫn theo yêu cầu, người rà soát phát hiện vấn đề mang tính '
       'hệ thống ở phần tài liệu tham khảo, nên đã tiến hành rà soát toàn bộ.')
P(doc, 'Phạm vi rà soát: (i) toàn bộ 21 tài liệu tham khảo trong danh mục; (ii) các luận điểm '
       'có số liệu hoặc mô tả nghiên cứu cụ thể trong thân bài. Phương pháp: đối chiếu từng tài '
       'liệu với PubMed, DOI, trang tạp chí gốc và bản PDF gốc (nếu có trong kho dự án).')
P(doc, 'Lưu ý: báo cáo này KHÔNG chỉnh sửa bản thảo. File build_bai2_KHGDVN.py được giữ '
       'nguyên để tác giả tự quyết định hướng xử lý.', italic=True)

# ---- 2. Ket qua tong hop ----
H(doc, '2. Kết quả tổng hợp', 1)
P(doc, 'Trên 21 tài liệu tham khảo:')
bullet(doc, '2 tài liệu trích dẫn CHÍNH XÁC (Beck 1976; Compas và cộng sự 2017).',
       bold_prefix='Đúng: ')
bullet(doc, '2 tài liệu CƠ BẢN CÓ THẬT nhưng cần chỉnh sửa chi tiết (luận án Trần Nguyễn '
            'Ngọc; báo cáo V-NAMHS của UNICEF).', bold_prefix='Tạm ổn: ')
bullet(doc, '17 tài liệu SAI nghiêm trọng: sai tên tác giả, sai tên bài, sai tạp chí, sai số '
            'tập/số bài — trong đó ít nhất 1 tài liệu (bài "Happy House" của "Trần T.M.T.") '
            'KHÔNG TỒN TẠI (bịa hoàn toàn).', bold_prefix='Sai / bịa: ', color=RED)
P(doc, 'Ngoài ra, thân bài có ít nhất 15 luận điểm sai sự thật hoặc mâu thuẫn nội bộ '
       '(xem Mục 4). Đáng lưu ý nhất: phần Phương pháp đang khẳng định "Mọi số liệu trích dẫn '
       'trong bài đều được đối chiếu trực tiếp với bài gốc" — khẳng định này hiện không đúng '
       'thực tế và là rủi ro liêm chính khoa học nếu bài được nộp.', color=RED)

# ---- 3. Bang ra soat TLTK ----
H(doc, '3. Bảng A — Rà soát 21 tài liệu tham khảo', 1)
P(doc, 'Cột "Trạng thái": ô nền xanh = đúng; nền vàng = tạm ổn, cần chỉnh; nền đỏ = sai/bịa.',
  italic=True, size=10)

ref_rows = [
 ('1', 'Beck, A. T. (1976). Cognitive therapy and the emotional disorders. International '
       'Universities Press.',
  'ĐÚNG',
  'Sách kinh điển, trích dẫn chính xác.'),
 ('2', 'Bradshaw, Lochman, Powell, Boxmeyer, Qu, Marchionno, Sterling (2025). "Effectiveness '
       'of the Early Adolescent Coping Power program: A cluster RCT." J School Psychology, '
       '108, 101321.',
  'SAI',
  'Bài có thật nhưng tác giả và tên bài sai. Bản đúng: Bradshaw C.P., McDaniel H.D., Pas '
  'E.T., Debnam K.J., Bottiani J.H., Powell N., Ialongo N.S., Morgan-Lopez A., Lochman J.E. '
  '(2025). "Randomized control trial of the Early Adolescent Coping Power Program: Effects '
  'on emotional and behavioral problems in middle schoolers." Journal of School Psychology. '
  'Cần xác minh lại số tập/số bài.'),
 ('3', 'Bress, Stewart, Fenster, Ehrlich, Magariños, Tirpak, Pine (2024). "A '
       'smartphone-delivered CBT intervention for anxiety symptoms in adolescents: A RCT." '
       'JAMA Network Open, 7(11), e2443572.',
  'SAI NẶNG',
  'Sai tác giả, tên bài, số issue, số bài. Bản đúng: Bress J.N., Falk A., Schier M.M., '
  'Jaywant A., Moroney E., Dargis M., Bennett S.M., Scult M.A., Volpp K.G., Asch D.A., '
  'Balachandran M., Perlis R.H., Lee F.S., Gunning F.M. (2024). "Efficacy of a Mobile '
  'App-Based Intervention for Young Adults With Anxiety Disorders: A Randomized Clinical '
  'Trial." JAMA Network Open, 7(8), e2428372. doi:10.1001/jamanetworkopen.2024.28372. '
  'Đối tượng là THANH NIÊN 18–25 tuổi.'),
 ('4', 'Brown-Carter, Patalay, Stallard (2025). "School-based emotional and behavioural '
       'interventions in the UK: A systematic review." J Mental Health, 34(2), 145–162.',
  'KHÔNG XÁC MINH ĐƯỢC',
  'Không tìm thấy bài đúng tên này. "Brown-Carter" không phải họ tác giả chuẩn — có thể là '
  'ghép nhầm một bài "Brown ... Carter" có thật (kho dự án có bài B5 "Brown Carter 2025" về '
  'trường học Anh). Cần đối chiếu file dự án 03_Ban-dich/B5 và xác minh lại toàn bộ.'),
 ('5', 'Cai, Wang, Liu, Chen (2025). "School-based resilience programmes and adolescent '
       'mental health: A systematic review and meta-analysis." Frontiers in Psychology, 16, '
       '1382764.',
  'SAI',
  'Bản đúng: Cai C., Mei Z., Wang Z., Luo S. (2025). "School-based interventions for '
  'resilience in children and adolescents: a systematic review and meta-analysis of '
  'randomized controlled trials." Frontiers in PSYCHIATRY (không phải Psychology), '
  'doi:10.3389/fpsyt.2025.1594658, PMID 40458775. Số bài đúng là 1594658. Nghiên cứu gốc: '
  '38 RCT / 15.730 người.'),
 ('6', 'Chen, Chen, Lin, Wang, Ho (2025). "A smartphone-delivered CBT for insomnia '
       'intervention and depression prevention in young adolescents: A RCT." PLOS Medicine, '
       '22(4), e1004465.',
  'SAI',
  'Bản đúng: "Effectiveness of app-based cognitive behavioral therapy for insomnia on '
  'preventing major depressive disorder in youth with insomnia and subclinical depression: '
  'A randomized clinical trial." PLOS Medicine, 21/01/2025, doi:10.1371/journal.pmed.1004510, '
  'PMID 39836656. Sai tên bài, số issue, số bài. Đối tượng 15–25 tuổi; mục tiêu phòng trầm '
  'cảm (không phải lo âu).'),
 ('7', 'Compas, Jaser, Bettis, Watson, Gruhn, Dunbar, Williams, Thigpen (2017). "Coping, '
       'emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis '
       'and narrative review." Psychological Bulletin, 143(9), 939–991.',
  'ĐÚNG',
  'Trích dẫn chính xác hoàn toàn. PMID 28616996. N = 80.850 / 212 nghiên cứu — khớp thân bài.'),
 ('8', 'He, X., Liu, Y., Zhou, J., Wang, Q. (2025). "Effectiveness of an 8-session CBT-based '
       'prevention program (Power Up – CBTD)..." J Affect Disord, 367, 412–421.',
  'SAI NẶNG',
  'Bản đúng: He Q., Li J., Wang J., Qu Z. (2026). "Preventing depression in Chinese children '
  'and adolescents: A pilot study of a brief school-based cognitive behavioral group '
  'program." Journal of Affective Disorders, 394, 120559. doi:10.1016/j.jad.2025.120559. '
  'Sai tác giả, tên bài, năm, số tập–trang.'),
 ('9', 'Li, S. H., Cheng, J. W., Yang, M., Newby, J. M., Christensen, H. (2024). '
       '"Co-designing a smartphone-based CBT intervention for adolescent depression: '
       'ClearlyMe." Cambridge Prisms: Global Mental Health, 11, e87.',
  'SAI',
  'Ứng dụng ClearlyMe có thật (của Úc — Black Dog Institute/UNSW, KHÔNG phải Đại học '
  'Cambridge). Bài đồng thiết kế thật: Li S.H. và cộng sự (2022). "A cognitive behavioural '
  'therapy smartphone app for adolescent depression and anxiety: co-design of ClearlyMe." '
  'The Cognitive Behaviour Therapist (không phải Cambridge Prisms: GMH), 2022. Sai tạp chí, '
  'năm, đồng tác giả.'),
 ('10', 'Liu, P., Han, H., Chen, J., Wang, Y. (2025). "Comparative efficacy of CBT delivery '
        'modes for adolescent generalized anxiety disorder: A network meta-analysis." '
        'Translational Psychiatry, 15, 124.',
  'SAI',
  'Bài thật: "CBT treatment delivery formats for generalized anxiety disorder: a systematic '
  'review and network meta-analysis of randomized controlled trials." Translational '
  'Psychiatry 2025, doi:10.1038/s41398-025-03414-3. Mẫu là NGƯỜI LỚN (52 thử nghiệm, 4.361 '
  'bệnh nhân, tuổi TB 43) — KHÔNG phải vị thành niên như tên trích dẫn. Tác giả chưa xác '
  'minh được.'),
 ('11', 'Matsumoto, Hamatani, Nagai, Sutoh, Nakagawa, Shimizu (2024). "Effectiveness of '
        'internet-based CBT on smartphone for adolescents with social anxiety disorder: RCT." '
        'J Medical Internet Research, 26, e51661.',
  'SAI',
  'Bản đúng: Matsumoto K., Hamatani S. và cộng sự (2024). "Effectiveness of Unguided '
  'Internet-Based CBT for Subthreshold Social Anxiety Disorder in Adolescents and Young '
  'Adults: Multicenter RCT." JMIR Pediatrics and Parenting (không phải J Medical Internet '
  'Research), e55786, PMID 39037759. Mẫu thật: 77 học sinh.'),
 ('12', 'Praptomojati, Suryani, Salim, Pratiwi (2024). "Culturally adapted CBT for adolescent '
        'anxiety in Southeast Asia: A scoping review." Asian Journal of Psychiatry, 92, '
        '103907.',
  'SAI MỘT PHẦN',
  'Bài thật: Praptomojati A. và cộng sự (2024). "A systematic review of Culturally Adapted '
  'Cognitive Behavioral Therapy (CA-CBT) for anxiety disorders in Southeast Asia." Asian '
  'Journal of Psychiatry, 02/2024, PMID 38199202. Là "systematic review" (không phải scoping '
  'review), cho rối loạn lo âu nói chung. Tác giả đầu, tạp chí, năm đúng; tên bài và đồng '
  'tác giả cần sửa.'),
 ('13', 'Qiaochu, Z., Wang, H., Zhang, M. (2025). "Mobile-delivered CBT for adolescent '
        'anxiety in China: A scoping review of efficacy and implementation." Clinical '
        'Psychology and Psychotherapy, 32(3), e3032.',
  'SAI',
  'Bài thật: "Effectiveness of Mobile-Based Cognitive Behavioural Therapy for Depression and '
  'Anxiety in Children and Young People: A Systematic Review of Randomized Controlled '
  'Trials." Clinical Psychology & Psychotherapy 2025, doi:10.1002/cpp.70173. Là tổng quan hệ '
  'thống RCT (không phải scoping review riêng Trung Quốc); số bài e3032 sai.'),
 ('14', 'Samele, Stallard, Verkuyl, Tudor (2025). "Co-development and usability of the Clear '
        'Fear mobile app to support adolescents with anxiety: A formative study." JMIR '
        'Formative Research, 9, e58921.',
  'SAI',
  'Ứng dụng Clear Fear (tổ chức stem4, Anh) có thật. Bài 2025 thật: "Evaluation of the Clear '
  'Fear Smartphone App for Young People Experiencing Anxiety: Uncontrolled Pre- and '
  'Post-Follow-Up Study." JMIR Formative Research 2025, vol 9, e55603. Sai tên bài và số bài '
  '(e58921); danh sách tác giả cần xác minh.'),
 ('15', 'Urao, Yoshida, Koshiba, Sato, Ishikawa, Shimizu (2018). "Effectiveness of a '
        'CBT-based anxiety prevention programme for children: A preliminary quasi-experimental '
        'study in Japan." Child Adolesc Psychiatry Ment Health, 12, 33.',
  'SAI (lai ghép)',
  'Tên bài bị lấy nhầm từ bản 2016. Bản 2018 đúng: "...anxiety prevention programme AT AN '
  'ELEMENTARY SCHOOL in Japan: a quasi-experimental study." Child and Adolescent Psychiatry '
  'and Mental Health, 12:33, doi:10.1186/s13034-018-0240-5. (Tên "for children: A '
  'preliminary..." thuộc bản 2016, doi:10.1186/s13034-016-0091-x.)'),
 ('16', 'Urao, Yoshinaga, Asano, Ishikawa, Tano, Sato, Shimizu (2022). "Effectiveness of a '
        'CBT-based anxiety prevention programme at an elementary school in Japan: A '
        'quasi-experimental study." Child Adolesc Psychiatry Ment Health, 16, 26.',
  'SAI NẶNG',
  'Bản đúng: Urao Y., Yoshida M., Sato Y., Shimizu E. (2022). "School-based cognitive '
  'behavioural intervention programme for addressing anxiety in 10- to 11-year-olds using '
  'short classroom activities in Japan: a quasi-experimental study." BMC PSYCHIATRY (không '
  'phải Child Adolesc Psychiatry Ment Health), 22:658, doi:10.1186/s12888-022-04326-y. '
  'Đồng tác giả sai.'),
 ('17', 'Walder, N., Berger, T., Krieger, T. (2025). "Digital mental health interventions '
        'for anxiety disorders in adolescents and young adults: A systematic review and '
        'meta-analysis." J Medical Internet Research, 27, e63278.',
  'SAI',
  'Bản đúng: Walder N., Frey A., Berger T., Schmidt S.J. (2025). "Digital Mental Health '
  'Interventions for the Prevention and Treatment of SOCIAL Anxiety Disorder in Children, '
  'Adolescents, and Young Adults: Systematic Review and Meta-Analysis of RCTs." JMIR 2025, '
  'e67067, PMID 40504611. Tác giả thứ 2–4 sai (không có "Krieger"); tên bài và số bài sai.'),
 ('18', 'Xian, J., Tian, M., Liu, X., Sun, Y. (2024). "Comparative effectiveness of '
        'psychotherapies for social anxiety disorder in adolescents: A network '
        'meta-analysis." Journal of Affective Disorders, 348, 222–233.',
  'SAI',
  'Bản đúng: Xian J., Zhang Y., Jiang B. (2024). "Psychological interventions for social '
  'anxiety disorder in children and adolescents: A systematic review and network '
  'meta-analysis." Journal of Affective Disorders, 365, 614–627, PMID 39173929. Đồng tác '
  'giả và số tập–trang sai.'),
 ('19', 'Trần Nguyễn Ngọc (2018). "Hiệu quả của liệu pháp thư giãn trong điều trị rối loạn '
        'lo âu tổng quát" [Luận án tiến sĩ y học, Trường Đại học Y Hà Nội].',
  'TẠM ỔN',
  'Luận án có thật (Trường Đại học Y Hà Nội — sdh.hmu.edu.vn). Tên chính thức gần đúng là '
  '"Đánh giá hiệu quả điều trị rối loạn lo âu lan TỎA bằng liệu pháp thư giãn – luyện tập". '
  'Nên sửa tên bài cho khớp bản gốc; xác nhận lại năm bảo vệ (2018/2019).'),
 ('20', 'Trần, T. M. T., Doan, V. H., Pham, A. T. (2023). "Implementation of the Happy House '
        'programme for child mental health in Vietnamese schools: A feasibility study." '
        'Cambridge Prisms: Global Mental Health, 10, e54.',
  'BỊA',
  'Không có bài này. Bài Happy House thật: Tran T.D., Nguyen H., Shochet I., Nguyen N., '
  'La N., Wurfl A., Orr J., Nguyen H., Stocker R., Fisher J. (2023). "School-based universal '
  'mental health promotion intervention for adolescents in Vietnam: Two-arm, parallel, '
  'controlled trial." Cambridge Prisms: Global Mental Health, 10, e69, '
  'doi:10.1017/gmh.2023.66, PMC10643236. Tác giả, tên bài, số bài (e54) đều bịa.'),
 ('21', 'UNICEF Vietnam (2022). "Vietnam National Adolescent Mental Health Survey (V-NAMHS) '
        'report." UNICEF Vietnam & Institute of Sociology.',
  'TẠM ỔN',
  'Khảo sát V-NAMHS có thật. Báo cáo "Viet Nam Adolescent Mental Health Survey (V-NAMHS) — '
  'Report on Main Findings" công bố 02/2023 (nên ghi 2023). Lưu ý: V-NAMHS là khảo sát hộ '
  'gia đình về tỷ lệ rối loạn tâm thần ở trẻ 10–17 tuổi — không khảo sát dịch vụ tâm lý học '
  'đường (xem lỗi thân bài B14).'),
]
audit_table(doc,
    ['STT', 'Trích dẫn trong bản thảo', 'Trạng thái', 'Vấn đề & bản đúng đã xác minh'],
    ref_rows, widths=[0.7, 5.0, 2.2, 7.6])

# to mau o trang thai
tblA = doc.tables[-1]
for ri, rd in enumerate(ref_rows):
    cell = tblA.rows[ri+1].cells[2]
    shade(cell, status_shade(rd[2]))

# ---- 4. Bang loi than bai ----
H(doc, '4. Bảng B — Lỗi luận điểm trong thân bài', 1)
P(doc, 'Ngoài tài liệu tham khảo, các luận điểm sau trong thân bài sai sự thật, mâu thuẫn '
       'nội bộ, hoặc trích sai nguồn:')

body_rows = [
 ('Mục 1; 3.1; 3.2', 'Chương trình "Happy House" "hợp tác / phối hợp với Đại học Cambridge".',
  'SAI. Happy House do Đại học Monash (Úc) + Đại học Y tế Công cộng Hà Nội + Đại học Công '
  'nghệ Queensland thực hiện. "Cambridge" chỉ là nhà xuất bản tạp chí (Cambridge Prisms: '
  'Global Mental Health), không phải đơn vị nghiên cứu.'),
 ('Mục 3.2', '"Happy House" thử nghiệm trên "trường tiểu học và trung học cơ sở".',
  'SAI. Thử nghiệm trên học sinh LỚP 10 (trung học phổ thông), 15–16 tuổi, tại 8 trường THPT '
  'ở Hà Nội.'),
 ('Mục 3.1', '"Happy House"... "can thiệp dựa trên CBT kéo dài 12 tuần", "hợp tác Việt – Anh".',
  'CẦN KIỂM. Happy House là bản Việt hoá chương trình RAP của Úc (hợp tác Việt – ÚC). Mốc '
  '"12 tuần" chưa khớp tài liệu gốc (đo lúc 2 tuần và 6 tháng sau can thiệp).'),
 ('Mục 3.2', 'Urao "Journey of the Brave" được "kiểm định qua hai thử nghiệm lâm sàng ngẫu '
             'nhiên (RCT)".',
  'SAI. Cả hai nghiên cứu Urao (2018 và 2022) đều là TỰA THỰC NGHIỆM (quasi-experimental), '
  'không phải RCT.'),
 ('Mục 3.2', '"Journey of the Brave – phiên bản CBT 14 buổi".',
  'CHƯA ĐỦ. Chỉ bản 2022 là 14 buổi × 20 phút; bản 2018 là 10 buổi × 45 phút.'),
 ('Mục 3.2', 'Power Up-CBTD là "can thiệp 8 buổi" cho "học sinh trung học cơ sở".',
  'CHƯA XÁC MINH / SAI MỘT PHẦN. Số buổi "8" chưa kiểm chứng được từ bài gốc. Mẫu thật gồm '
  'học sinh lớp 5, 6 và 9 (có cả bậc tiểu học), tại tỉnh Hà Nam, Trung Quốc.'),
 ('Mục 3.2', 'Power Up-CBTD "kết quả cho thấy cải thiện đáng kể trên cả triệu chứng lo âu" — '
             'làm "bằng chứng cho can thiệp transdiagnostic".',
  'SAI. Theo bài gốc (He và cộng sự), lo âu có giảm nhưng KHÔNG đạt ý nghĩa thống kê (chênh '
  'lệch giữa nhóm AMD = −6,68; p = 0,07). Luận điểm "transdiagnostic" đang dựa trên dữ liệu '
  'sai.'),
 ('Mục 3.3', 'Ứng dụng "Maya" "dành cho vị thành niên Mỹ", "hơn 200 vị thành niên", so với '
             '"nhóm giả dược ứng dụng (sham app)".',
  'SAI cả ba chi tiết. Đối tượng là THANH NIÊN 18–25 tuổi; mẫu n = 59; thiết kế so 3 nhánh '
  'khích lệ, KHÔNG có nhóm giả dược / không can thiệp.'),
 ('Mục 3.3', 'Matsumoto thử nghiệm "trên mẫu vài trăm học sinh".',
  'SAI. Mẫu thật 77 người (38 can thiệp + 39 đối chứng).'),
 ('Mục 3.3', 'ClearlyMe (Li và cộng sự) "hợp tác với Đại học Cambridge".',
  'SAI. ClearlyMe là ứng dụng của Úc (Black Dog Institute / UNSW).'),
 ('Mục 3.3', 'Qiaochu "tổng quan các can thiệp CBT trên điện thoại tại Trung Quốc xác định '
             '12 ứng dụng... ba ứng dụng đã hoàn thành thử nghiệm pha 3".',
  'NGHI BỊA. Bài gốc là tổng quan hệ thống RCT về CBT qua di động cho trẻ em / thanh thiếu '
  'niên nói chung — không phải bản kiểm kê "12 ứng dụng tại Trung Quốc".'),
 ('Mục 3.2', 'Cai "trên 39 thử nghiệm".',
  'SAI. Nghiên cứu gốc: 38 RCT (21 đưa vào phân tích gộp).'),
 ('Mục 3.1 vs Mục 1', 'Cùng bài Liu (2025): Mục 3.1 ghi "47 thử nghiệm" + g = 0,82/0,68/0,55; '
             'Mục 1 ghi "52 thử nghiệm" + SMD = 1,62.',
  'MÂU THUẪN NỘI BỘ. Bài gốc có 52 thử nghiệm. Mục 3.1 còn gọi mẫu là "vị thành niên" trong '
  'khi bài gốc là người lớn (tuổi TB 43).'),
 ('Mục 3.3', 'Báo cáo UNICEF Việt Nam (2022): "khoảng 60% trường được khảo sát có ít nhất '
             'một hoạt động hỗ trợ tâm lý học đường".',
  'SAI NGUỒN. V-NAMHS là khảo sát tỷ lệ rối loạn tâm thần ở trẻ 10–17 tuổi, không khảo sát '
  'dịch vụ tâm lý trường học; số "60%" không có trong V-NAMHS.'),
 ('Mục 2 (Phương pháp)', '"Mọi số liệu trích dẫn trong bài đều được đối chiếu trực tiếp với '
             'bài gốc trước khi đưa vào bản thảo."',
  'KHÔNG ĐÚNG THỰC TẾ. Rà soát cho thấy phần lớn tài liệu chưa được đối chiếu. Câu này cần '
  'xoá hoặc viết lại trung thực — nếu giữ nguyên sẽ là rủi ro liêm chính khoa học.'),
]
audit_table(doc,
    ['Vị trí', 'Bản thảo đang ghi', 'Đánh giá & sự thật đã xác minh'],
    body_rows, widths=[2.3, 6.3, 6.9])

tblB = doc.tables[-1]
for ri, rd in enumerate(body_rows):
    cell = tblB.rows[ri+1].cells[2]
    s = rd[2].split('.')[0].upper()
    if 'MÂU THUẪN' in s or s.startswith('SAI') or 'KHÔNG ĐÚNG' in s or 'NGHI BỊA' in s:
        shade(cell, RED_SH)
    else:
        shade(cell, YELLOW_SH)

# ---- 5. Khuyen nghi ----
H(doc, '5. Khuyến nghị', 1)
bullet(doc, 'Không nộp bản thảo ở trạng thái hiện tại. Tỷ lệ lỗi tài liệu tham khảo quá cao '
            '(17/21) khiến bài không đạt chuẩn liêm chính trích dẫn.',
       bold_prefix='1. ')
bullet(doc, 'Dựng lại danh mục tài liệu tham khảo từ đầu: với mỗi tài liệu, tra trực tiếp '
            'PubMed/DOI, lấy đúng tác giả – tên bài – tạp chí – tập/số – trang/số bài, và lưu '
            'mã DOI/PMID. Báo cáo này (Bảng A) đã cung cấp sẵn bản đúng cho phần lớn các bài.',
       bold_prefix='2. ')
bullet(doc, 'Sửa lại từng luận điểm thân bài theo Bảng B, đặc biệt: bỏ khẳng định "cải thiện '
            'đáng kể trên lo âu" của Power Up-CBTD; sửa mô tả ứng dụng Maya; thống nhất số '
            'liệu bài Liu giữa Mục 1 và Mục 3.1.', bold_prefix='3. ')
bullet(doc, 'Viết lại câu trong phần Phương pháp về việc "đối chiếu trực tiếp với bài gốc" '
            'cho đúng quy trình đã thực sự thực hiện.', bold_prefix='4. ')
bullet(doc, 'Sau khi tác giả chốt nội dung sửa, người rà soát có thể cập nhật trực tiếp file '
            'build_bai2_KHGDVN.py và chạy lại để sinh bản .docx mới. Báo cáo này không tự ý '
            'sửa bản thảo.', bold_prefix='5. ')

# ---- 6. Tham khao ----
H(doc, '6. Tham khảo & truy vết', 1)
P(doc, 'Nguồn gốc đã đối chiếu để xác minh (các bài cốt lõi):', bold=True, size=11)
ref_item(doc,
    'Bress J.N., Falk A., Schier M.M., và cộng sự (2024). "Efficacy of a Mobile App-Based '
    'Intervention for Young Adults With Anxiety Disorders: A RCT." JAMA Network Open, '
    '7(8):e2428372.',
    'doi:10.1001/jamanetworkopen.2024.28372',
    'https://doi.org/10.1001/jamanetworkopen.2024.28372')
ref_item(doc,
    'He Q., Li J., Wang J., Qu Z. (2026). "Preventing depression in Chinese children and '
    'adolescents: A pilot study of a brief school-based cognitive behavioral group program." '
    'Journal of Affective Disorders, 394:120559.',
    'doi:10.1016/j.jad.2025.120559',
    'https://doi.org/10.1016/j.jad.2025.120559')
ref_item(doc,
    'Tran T.D., Nguyen H., Shochet I., và cộng sự (2023). "School-based universal mental '
    'health promotion intervention for adolescents in Vietnam." Cambridge Prisms: Global '
    'Mental Health, 10:e69. PMCID: PMC10643236.',
    'doi:10.1017/gmh.2023.66',
    'https://doi.org/10.1017/gmh.2023.66')
ref_item(doc,
    'Urao Y. và cộng sự (2018). Child and Adolescent Psychiatry and Mental Health, 12:33.',
    'doi:10.1186/s13034-018-0240-5', 'https://doi.org/10.1186/s13034-018-0240-5')
ref_item(doc,
    'Urao Y. và cộng sự (2022). BMC Psychiatry, 22:658.',
    'doi:10.1186/s12888-022-04326-y', 'https://doi.org/10.1186/s12888-022-04326-y')
ref_item(doc,
    'Cai C., Mei Z., Wang Z., Luo S. (2025). Frontiers in Psychiatry. PMID 40458775.',
    'doi:10.3389/fpsyt.2025.1594658', 'https://doi.org/10.3389/fpsyt.2025.1594658')
ref_item(doc,
    'Walder N., Frey A., Berger T., Schmidt S.J. (2025). JMIR, 27:e67067. PMID 40504611.',
    'doi:10.2196/67067', 'https://doi.org/10.2196/67067')
ref_item(doc,
    'Xian J., Zhang Y., Jiang B. (2024). Journal of Affective Disorders, 365:614–627. '
    'PMID 39173929.')
ref_item(doc,
    'Compas B.E. và cộng sự (2017). Psychological Bulletin, 143(9):939–991. PMID 28616996 — '
    'tài liệu trong bản thảo trích dẫn ĐÚNG.')

P(doc, 'Truy vết nội bộ:', bold=True, size=11)
bullet(doc, '06_Scripts/build_bai2_KHGDVN.py — bản thảo được rà soát; danh mục tài liệu tham '
            'khảo nằm ở biến TLTK (dòng 169–193), thân bài ở các biến SECTION_* '
            '(dòng 59–166).', size=11)
bullet(doc, '02_Papers-goc/QT073_He_2025_PowerUpCBTD_China_JAD.pdf (trang 1) — đối chiếu bài '
            'He và cộng sự.', size=11)
bullet(doc, '03_Ban-dich/B3_JAMA_App_raw.txt — văn bản gốc nghiên cứu ứng dụng Maya (Bress và '
            'cộng sự 2024).', size=11)
bullet(doc, '06_Scripts/dich_VN030_HappyHouse.py — bản dịch bài Happy House (VN030). Lưu ý: '
            'trong phiên làm việc này đã sửa số định danh bài Happy House trong file đó từ '
            '"e58" thành "e69" (dòng 102).', size=11)
bullet(doc, '02_Papers-goc/canonical_index.json — các mã QT068, QT069, QT073.', size=11)

P(doc, 'Ghi chú: báo cáo này chỉ rà soát và đề xuất; file build_bai2_KHGDVN.py KHÔNG bị '
       'chỉnh sửa.', italic=True, size=11)

doc.save(OUT)
print('DA LUU:', OUT)
print('So bang:', len(doc.tables), '| So doan:', len(doc.paragraphs))
