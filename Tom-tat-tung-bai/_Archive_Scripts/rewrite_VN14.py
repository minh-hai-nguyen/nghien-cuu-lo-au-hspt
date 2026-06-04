# -*- coding: utf-8 -*-
"""VN14: Hoàng Trung Học 2025 — viết lại chuẩn CTH v5 (như bài 02-11)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.5
for s in doc.sections: s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph(); r = p.add_run(text); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(text, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BƯỚC 1: ĐỊNH DANH =====
bl('Tóm tắt bài VN-14', bold=True)
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình nghiên cứu \u00ab Mức độ căng thẳng, lo âu và trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 tại Việt Nam: Nghiên cứu cắt ngang \u00bb (Levels of Stress, Anxiety, and Depression in Adolescents during and after the COVID-19 Pandemic in Vietnam: A Cross-sectional Study), do Hoàng Trung Học (Tiến sĩ, Học viện Quản lý Giáo dục Quốc gia) và Nguyễn Thùy Dung (Thạc sĩ, Viện Nghiên cứu, Đào tạo và Ứng dụng Tâm lý) thực hiện (2025), khảo sát trên mẫu 8.473 thanh thiếu niên từ lớp 6 đến lớp 12 tại 6 tỉnh/thành phố đại diện miền Nam (TPHCM, Đồng Nai, Long An) và miền Bắc (Hà Nội, Hải Dương, Vĩnh Phúc), Việt Nam, xuất bản trên American Journal of Psychiatric Rehabilitation, tập 28, số 1, tháng 4/2025, trang 360\u2013367.')

# ===== BƯỚC 2: TỔNG QUAN PP + KT1 =====
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21 \u2014 Depression Anxiety Stress Scale; Lovibond & Lovibond, 1995) phiên bản tiếng Việt, kết hợp với bảng hỏi bổ sung về giấc ngủ, sử dụng thiết bị điện tử, hoạt động thể chất và quan hệ gia đình. Nói cách khác, công trình này sử dụng phương pháp sàng lọc cắt ngang tại hai thời điểm \u2014 không phải nghiên cứu dọc thực sự \u2014 để so sánh mức độ rối loạn cảm xúc trước và sau khi COVID-19 kết thúc.', bold=True)

# ===== BƯỚC 3: ĐỊNH NGHĨA KT5 =====
bl('DASS-21 là phiên bản rút gọn 21 mục của Thang đo Trầm cảm Lo âu Căng thẳng gốc 42 mục (Lovibond & Lovibond, 1995), sàng lọc đồng thời ba tình trạng: trầm cảm, lo âu và căng thẳng. Phiên bản tiếng Việt đã được xác thực với Cronbach alpha > 0,7 cho tất cả thang con (Nunnally & Bernstein, 1994).')

# ===== BƯỚC 4: BIỆN MINH KT4 =====
bl('Tổng quan tài liệu của chúng tôi cho thấy trước COVID-19, tỷ lệ lo âu ở thanh thiếu niên Việt Nam dao động 12\u201315%, nhưng trong đại dịch tăng mạnh lên 37\u201347% (Nguyen và cộng sự, 2021). Tuy nhiên, dữ liệu so sánh trong vs sau COVID tại Việt Nam còn rất hạn chế. Đây là nghiên cứu có cỡ mẫu lớn nhất (n = 8.473) thực hiện so sánh này tại 6 tỉnh đại diện.', bold=True)

# ===== BƯỚC 5: LIỆT KÊ DỮ LIỆU =====
bl('Dữ liệu nhân khẩu xã hội bao gồm giới tính, khu vực cư trú (nông thôn/thành thị), khối lớp (6\u201312). Các yếu tố hành vi: thời gian ngủ ban ngày, thời gian sử dụng máy tính và điện thoại, thời gian thể thao, tính chất mối quan hệ với cha mẹ, hình thức học tập (trực tuyến/trực tiếp), mức độ giãn cách xã hội.')

# ===== BƯỚC 6: MÔ TẢ QUY TRÌNH =====
bl('Nghiên cứu cắt ngang tại hai thời điểm: Đợt 1 \u2014 tháng 12/2021 (trong COVID, n = 4.052), Đợt 2 \u2014 tháng 9/2023 (sau COVID, n = 4.337). Lấy mẫu thuận tiện ngẫu nhiên (convenient random sampling) từ các trường THCS và THPT tại 6 tỉnh. Nông thôn chiếm 62,8% mẫu. Phân tích: SPSS, thống kê mô tả, tương quan Pearson, hồi quy đa biến. Kiểm soát đa cộng tuyến (VIF < 2) và tự tương quan (DW = 0,353 và 0,340).')

# ===== BƯỚC 9: ĐỊNH LƯỢNG — BẢNG =====
rb('Kết quả nghiên cứu định lượng')

bl('Tỷ lệ căng thẳng, lo âu và trầm cảm trong vs sau COVID-19 (N = 8.389):', bold=True)
tbl(
    ['Tình trạng', 'Trong COVID (%)', 'Sau COVID (%)', 'Thay đổi'],
    [['Căng thẳng tổng', '65,5%', '55,4%', '\u221910,1 điểm'],
     ['  Nặng + Rất nặng', '33,0%', '18,6%', '\u221914,4 điểm'],
     ['Lo âu tổng', '41,5%', '25,4%', '\u221916,1 điểm'],
     ['  Nặng + Rất nặng', '14,5%', '7,2%', '\u22197,3 điểm'],
     ['Trầm cảm tổng', '34,2%', '20,1%', '\u221914,1 điểm'],
     ['  Nặng + Rất nặng', '8,1%', '4,0%', '\u22194,1 điểm'],
     ['Điểm DASS-21 TB', '26,68', '22,07', '\u22194,61 (p<0,01)']],
    widths=[4.0, 3.0, 3.0, 2.5])
doc.add_paragraph()

bl('Yếu tố ảnh hưởng \u2014 hồi quy đa biến (trong COVID, R\u00b2 = 0,190):', bold=True)
tbl(
    ['Yếu tố', 'Beta', 'Sig.', 'Ý nghĩa'],
    [['Quan hệ cha mẹ-con', '0,272', '0,000', 'MẠNH NHẤT \u2014 gia đình bất hòa tăng DAS'],
     ['Thời gian dùng điện tử', '0,176', '0,000', 'Dùng nhiều → tăng DAS'],
     ['Giấc ngủ ban ngày', '\u20130,149', '0,000', 'Ngủ đủ → giảm DAS (bảo vệ)'],
     ['Thời gian thể thao', '\u20130,087', '0,000', 'Thể thao → giảm DAS (bảo vệ)'],
     ['Tuổi', '\u20130,086', '0,000', 'Trẻ hơn → DAS cao hơn'],
     ['Hình thức học tập', '0,074', '0,000', 'Trực tuyến → tăng DAS'],
     ['Giãn cách xã hội', '0,073', '0,000', 'Giãn cách nhiều → tăng DAS'],
     ['Giới tính', '0,053', '0,001', 'Ảnh hưởng nhẹ']],
    widths=[4.0, 1.5, 1.5, 5.5])
doc.add_paragraph()

# ===== BƯỚC 10: ĐỐI CHIẾU KT3 =====
bh3('Khác biệt giai đoạn : Y văn quốc tế luôn báo cáo sự gia tăng các vấn đề SKTT ở thanh thiếu niên trong đại dịch COVID-19 (Racine và cộng sự, 2021). Thực tế, Hoàng Trung Học (2025) xác nhận xu hướng này: lo âu trong COVID 41,5% giảm xuống 25,4% sau COVID. Tuy nhiên, tỷ lệ sau COVID (25,4%) vẫn cao hơn trước dịch (12\u201315% theo Do & Nguyen, 2020) \u2014 gợi ý tác động kéo dài của đại dịch.')

# ===== BƯỚC 11+12: NHẬN XÉT =====
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về vai trò gia đình, công nghệ và giấc ngủ')
bl('*Quan hệ cha mẹ-con \u2014 yếu tố mạnh nhất.* Nổi bật nhất là mối quan hệ cha mẹ-con (Beta = 0,272) là yếu tố ảnh hưởng mạnh nhất đến mức độ DAS cả trong và sau COVID. Phát hiện này phù hợp với Pham và cộng sự (2024) tại Huế: chăm sóc cảm xúc giảm lo âu mạnh (beta = \u20130,40), và Qiu và cộng sự (2022) tại Trung Quốc: nuôi dạy tích cực giảm 70% trầm cảm (OR = 0,30).')
bl('*Sử dụng thiết bị điện tử.* Thời gian dùng máy tính/điện thoại (Beta = 0,176) là yếu tố nguy cơ mạnh thứ hai, tương tự Chen và cộng sự (2023) tại Trung Quốc (rối loạn chơi game OR = 5,00) và Nature Human Behaviour (2025): VTN có rối loạn SKTT dùng mạng xã hội nhiều hơn.')
bl('*Giấc ngủ và thể thao.* Giấc ngủ đủ (Beta = \u20130,149) và hoạt động thể thao (Beta = \u20130,087) là yếu tố bảo vệ, phù hợp với Zhu và cộng sự (2025): ngủ <5h AOR = 13,71 cho trầm cảm, hoạt động ngoài trời 2\u20133h AOR = 0,557 bảo vệ.')

# ===== BƯỚC 13: KẾT LUẬN (Kiến trúc B) =====
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy sự giảm đáng kể nhưng vẫn ở mức cao của triệu chứng lo âu (từ 41,5% xuống 25,4%) ở 8.473 thanh thiếu niên 6 tỉnh Việt Nam sau COVID-19 \u2014 cỡ mẫu lớn nhất trong các nghiên cứu SKTT VTN tại Việt Nam, gợi ý rằng thanh thiếu niên có khả năng phục hồi nhất định nhưng tỷ lệ hậu COVID vẫn cao hơn trước dịch. Ba yếu tố có thể can thiệp được \u2014 quan hệ gia đình (Beta = 0,272), sử dụng điện tử (Beta = 0,176) và giấc ngủ (Beta = \u20130,149) \u2014 là mục tiêu ưu tiên cho chương trình phòng ngừa. Trong bối cảnh hậu COVID tại Việt Nam, nhu cầu cải thiện quan hệ gia đình, quản lý thời gian sử dụng công nghệ và duy trì giấc ngủ chất lượng là thiết yếu.', bold=True)

# ===== PHẢN BIỆN =====
rh2('Phản biện')
bl('Cỡ mẫu lớn nhất VN (8.473), 6 tỉnh đại diện Bắc-Nam, hồi quy đa biến kiểm soát VIF và DW. Tuy nhiên, tạp chí AJPR không có impact factor rõ ràng, không lập chỉ mục PubMed/Scopus \u2014 giảm độ tin cậy học thuật. Hai đợt khảo sát trên 2 MẪU KHÁC NHAU (4.052 vs 4.337) \u2014 không phải nghiên cứu dọc thực sự, không thể theo dõi thay đổi ở cùng cá nhân. DASS-21 là sàng lọc \u2014 so với V-NAMHS 2022 (DISC-5: 2,3%), chênh lệch rất lớn. R\u00b2 = 0,190 \u2014 mô hình chỉ giải thích 19% phương sai. "Convenient random sampling" \u2014 mâu thuẫn (thuận tiện \u2260 ngẫu nhiên). Không phân tích giới tính chi tiết cho tỷ lệ lo âu.')

rh2('Hướng nghiên cứu tiếp theo')
bl('Nghiên cứu dọc thực sự theo dõi cùng nhóm HS từ trong COVID đến 2\u20133 năm sau. So sánh DASS-21 với DISC-5 trên cùng mẫu VTN. Phân tích giới tính chi tiết. Đánh giá hiệu quả can thiệp: đào tạo kỹ năng cha mẹ, quản lý thời gian điện tử tại trường. Xuất bản trên tạp chí có chỉ mục quốc tế.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50 Trung bình-Khá. Cỡ mẫu lớn nhất VN (8.473), so sánh COVID vs hậu COVID, hồi quy đa biến, nhưng tạp chí yếu và không phải NC dọc thực sự.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('VN14_HoangTrungHoc_2025.docx')
sys.stderr.write('VN14_HoangTrungHoc_2025.docx OK\n')
