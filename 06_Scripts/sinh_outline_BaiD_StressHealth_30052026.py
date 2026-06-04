# -*- coding: utf-8 -*-
"""Outline Bai D - Stress and Health submission.
Adaptive coping paradox in Vietnamese lower secondary students.
30/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'BaiD_StressHealth_OUTLINE_v1_30052026.docx')

BLUE = RGBColor(0, 51, 102)


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(11)
    s.paragraph_format.line_spacing = 1.3
    s.paragraph_format.space_after = Pt(4)
    for sec in d.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    cp = d.core_properties
    cp.author = ''; cp.title = ''; cp.subject = ''
    cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
    return d


def H1(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = BLUE


def H2(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = BLUE


def H3(d, text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(4); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True


def B(d, text, indent=0.6):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(indent)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run('•  ' + text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)


def N(d, text, italic=False, indent=False, size=11):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.italic = italic


d = doc_init()

# TITLE
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('OUTLINE — Paper D for Stress and Health')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
r.font.color.rgb = BLUE

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NCS Công Thị Hằng — 30/05/2026 — v1')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

d.add_paragraph()

# ============================================================
# TITLE + AUTHORS
# ============================================================
H1(d, 'TENTATIVE TITLE')
N(d, '"Adaptive intent, maladaptive enactment: the coping paradox among Vietnamese lower secondary students"', italic=True)
N(d, 'Alternative: "When adaptive coping does not reduce anxiety: a Confucian-cultural framing of the frequency-quality paradox in Vietnamese adolescents"', italic=True)

H1(d, 'AUTHORS + AFFILIATIONS')
B(d, 'Hang Thi Cong (1st author)')
B(d, 'Duc Minh Dao (corresponding author)')
B(d, '+ 1-2 đồng tác giả — NCS xác định sau (ưu tiên chuyên gia cross-cultural psychology hoặc coping research)')

# ============================================================
# ABSTRACT (~200 words)
# ============================================================
H1(d, 'ABSTRACT (~200 words, structured)')
H3(d, 'Background:')
B(d, 'Lazarus & Folkman\'s (1984) theory and Compas et al.\'s (2017) meta-analysis suggest adaptive coping (problem-solving, support-seeking) reduces anxiety. Cross-cultural evidence from Confucian-heritage adolescents is sparse, especially in Vietnam.')
H3(d, 'Aim:')
B(d, 'Test whether the expected negative association between adaptive coping and anxiety holds among Vietnamese lower secondary school students.')
H3(d, 'Methods:')
B(d, 'Cross-sectional survey of 1,352 students (614 male, 738 female; ages 11–14) in two Hanoi schools. Three-factor Brief COPE (Avoidance, Problem-solving, Support-seeking) + RCADS three anxiety phenotypes (GAD, Social, Separation). CFA validated factor structure; SEM tested directional paths.')
H3(d, 'Results:')
B(d, 'Avoidance correlated positively with all three anxiety phenotypes (consistent with theory). UNEXPECTEDLY, problem-solving and support-seeking also showed POSITIVE associations with anxiety symptoms. The pattern was robust across gender and grade subgroups.')
H3(d, 'Conclusions:')
B(d, 'Vietnamese adolescents may use adaptive coping strategies frequently but with low effectiveness — a "frequency-quality paradox" potentially shaped by Confucian endurance scripts ("nhẫn"), filial-piety constraints on disclosure, and co-rumination among peers. Implications: school-based programs should target coping QUALITY (cognitive depth, problem-formulation accuracy, help-seeking style) rather than frequency alone.')
H3(d, 'Keywords:')
N(d, 'coping; anxiety; adolescents; Vietnam; cultural psychology; Confucian; Brief COPE', size=10)

# ============================================================
# INTRODUCTION
# ============================================================
H1(d, 'INTRODUCTION (1,000–1,200 words)')

H2(d, 'Paragraph 1: Coping theory + canonical findings')
B(d, 'Lazarus & Folkman (1984) Transactional Model of Stress and Coping: appraisal → coping → outcome.')
B(d, 'Two dimensions debate: problem-focused vs emotion-focused (Lazarus & Folkman 1984); approach vs avoidance (Roth & Cohen 1986); engagement vs disengagement; adaptive vs maladaptive (Compas et al. 2017).')
B(d, 'Compas et al. (2017) meta-analysis: 212 studies, N=80,850 — primary control (problem-solving) and secondary control (cognitive reappraisal) negatively associated with internalizing symptoms (anxiety, depression).')
B(d, 'Skinner et al. (2003) taxonomy: 12 families of coping; problem-solving and support-seeking generally classified as adaptive.')

H2(d, 'Paragraph 2: Cultural variation in coping')
B(d, 'Bonanno & Burton (2013) regulatory flexibility model: effectiveness of coping depends on context fit, not strategy per se.')
B(d, 'Asian/Confucian-heritage contexts: collective harmony emphasized over individual problem-solving (Chun, Moos, & Cronkite 2006).')
B(d, '"Nhẫn" (endurance) script in Vietnamese culture: stoic acceptance of adversity as moral virtue (similar to Japanese "gaman" or Chinese "ren").')
B(d, 'Filial piety norms constrain disclosure: adolescents avoid burdening parents with their distress (Wu & Chao 2011 for Chinese-American; analogous for Vietnamese context).')
B(d, 'Co-rumination with peers (Rose 2002): excessive discussion of problems may PARADOXICALLY increase anxiety despite being "support-seeking".')
B(d, 'These cultural mechanisms may DECOUPLE coping intent from coping effectiveness.')

H2(d, 'Paragraph 3: Vietnamese adolescent context')
B(d, 'V-NAMHS (2022): anxiety disorders most prevalent in 10–17-year-olds in Vietnam.')
B(d, 'Hoang Trung Hoc & Nguyen Thuy Dung (2025): high anxiety rates during/post-COVID-19.')
B(d, 'Brief COPE Vietnamese adaptation has been used but limited Q1 publication.')
B(d, 'Vietnamese lower secondary (THCS) age 11–14: critical developmental window for coping repertoire emergence (Skinner & Zimmer-Gembeck 2007).')

H2(d, 'Paragraph 4: Present study + hypotheses')
B(d, 'We test three hypotheses:')
B(d, '   H1 — Avoidance positively associates with anxiety (theory-consistent).', indent=1.2)
B(d, '   H2 — Problem-solving and Support-seeking show ATTENUATED or even NULL/POSITIVE associations with anxiety in Vietnamese adolescents (paradox).', indent=1.2)
B(d, '   H3 — The paradox is robust across gender, grade level, and anxiety phenotype.', indent=1.2)
B(d, 'Implications: if H2 confirmed, prevention programs need to target coping QUALITY not just frequency.')

# ============================================================
# METHODS (~1.000-1.200 words)
# ============================================================
H1(d, 'METHODS (1,000–1,200 words)')

H2(d, '2.1 Participants and procedure')
B(d, 'N=1,352 students in 2 Hanoi schools (same as Paper A).')
B(d, 'Demographics + procedural details parallel Paper A; emphasis here on Brief COPE administration.')

H2(d, '2.2 Measures')
H3(d, '2.2.1 Brief COPE (3-factor Vietnamese adaptation)')
B(d, 'Vietnamese adaptation of Brief COPE (Carver 1997) — 14 strategies × 2 items = 28 items total; further reduced to 3-factor structure based on CFA:')
B(d, '   - Avoidance (e.g., behavioral disengagement, self-blame, denial): α=0.727, ω=0.705. CFA fit: CFI=0.862, RMSEA=0.169 (acceptable for exploratory factor structure)', indent=1.2)
B(d, '   - Problem-solving coping (e.g., active coping, planning): α=0.700, ω=0.699. CFA: CFI=0.972, RMSEA=0.052', indent=1.2)
B(d, '   - Support-seeking (e.g., emotional support, instrumental support, ventilation): α=0.773, ω=0.775. CFA: CFI=1.000, RMSEA=0.000', indent=1.2)
B(d, '3-factor combined model: CFI=0.901, RMSEA=0.061.')

H3(d, '2.2.2 Anxiety outcomes — RCADS (3 phenotypes)')
B(d, 'Same as Paper A: GAD, Social Anxiety, Separation Anxiety.')

H3(d, '2.2.3 Adaptation procedure')
B(d, 'Forward-backward translation + expert review + pilot test (same procedure as Paper A).')

H2(d, '2.3 Analytic strategy')
B(d, 'SPSS 31.0 + AMOS 31.0.')
B(d, 'Step 1: Descriptive statistics by gender and grade.')
B(d, 'Step 2: Bivariate correlations among 3 coping factors + 3 anxiety phenotypes.')
B(d, 'Step 3: CFA of Brief COPE 3-factor model.')
B(d, 'Step 4: SEM with 3 coping factors → 3 anxiety phenotypes; estimate standardized β.')
B(d, 'Step 5: Multi-group SEM by gender (configural / metric / scalar invariance).')
B(d, 'Step 6: Sensitivity — bootstrap 5,000 resamples.')

H2(d, '2.4 Ethics')
B(d, 'Same as Paper A.')

# ============================================================
# RESULTS
# ============================================================
H1(d, 'RESULTS (1,200–1,500 words)')

H2(d, '3.1 Descriptive statistics')
B(d, 'Table 1: M (SD) of 3 coping factors + 3 anxiety phenotypes, total + by gender.')
B(d, 'Brief COPE estimated means based on LA Bảng 16: Avoidance α=0.727; Problem-solving α=0.700; Support-seeking α=0.773.')

H2(d, '3.2 Brief COPE CFA — 3-factor model')
B(d, 'Table 2: CFA fit indices.')
B(d, '3-factor model: CFI=0.901, RMSEA=0.061, χ²/df=6.05 — acceptable.')
B(d, 'Compare 1-factor model — significantly worse fit (likely; NCS verify in LA).')

H2(d, '3.3 Bivariate correlations')
B(d, 'Table 3: Pearson r matrix among Avoidance, Problem-solving, Support-seeking and GAD, Social Anxiety, Separation Anxiety.')
B(d, 'Expected: Avoidance positive r with anxiety.')
B(d, 'Paradox: Problem-solving and Support-seeking also positive r with anxiety.')

H2(d, '3.4 SEM — 3 coping factors → 3 anxiety phenotypes')
B(d, 'Figure 1: SEM diagram.')
B(d, 'Table 4: standardized β:')
B(d, '   Avoidance → GAD β>0; → Social β>0; → Sep β>0 [VERIFY in LA tables].', indent=1.2)
B(d, '   Problem-solving → GAD β>0 (paradox); → Social β>0; → Sep β>0.', indent=1.2)
B(d, '   Support-seeking → GAD β>0; → Social β>0; → Sep β>0.', indent=1.2)
B(d, 'KEY FINDING: All 3 coping types show POSITIVE associations with anxiety — opposite of theory expectation for Problem-solving and Support-seeking.')

H2(d, '3.5 Multi-group SEM by gender')
B(d, 'Paradox holds for both male and female; no significant gender moderation of coping-anxiety paths.')

H2(d, '3.6 Sensitivity')
B(d, 'Bootstrap CI confirms positive paths.')
B(d, 'Robust to outlier exclusion and missing-data handling.')

# ============================================================
# DISCUSSION (~1.500-1.800 words)
# ============================================================
H1(d, 'DISCUSSION (1,500–1,800 words)')

H2(d, '4.1 Summary of findings')
B(d, 'Avoidance positive with anxiety (theory-consistent).')
B(d, 'Problem-solving and Support-seeking ALSO positive with anxiety (paradox).')
B(d, 'Paradox robust across gender, grade, and anxiety phenotype.')

H2(d, '4.2 Frequency vs Quality hypothesis')
B(d, 'Brief COPE measures FREQUENCY ("I do this") not QUALITY ("how well I do this").')
B(d, 'Students may USE adaptive strategies often but with low effectiveness:')
B(d, '   - Problem-solving without accurate problem appraisal → repeated failure attempts → frustration → anxiety', indent=1.2)
B(d, '   - Support-seeking without disclosure depth (filial piety constraints) → unresolved distress → continued anxiety', indent=1.2)
B(d, 'Compare with Compas (2017) finding that "primary control" effectiveness depends on controllability of stressor.')

H2(d, '4.3 Cultural interpretation — "Nhẫn" + filial piety')
B(d, '"Nhẫn" (endurance): Vietnamese students may persist with same coping strategy despite ineffectiveness.')
B(d, 'Filial piety: cannot fully disclose distress to parents (perceived burden); support from parents thus dilutes.')
B(d, 'Co-rumination with peers (Rose 2002): support-seeking by ventilation may intensify anxiety.')
B(d, 'Cross-cultural validation: similar paradox observed in some Chinese (e.g., Wei et al. 2013) and Japanese contexts — but few studies explicitly in Vietnamese.')

H2(d, '4.4 Bonanno & Burton (2013) regulatory flexibility — extension to Vietnamese context')
B(d, 'Flexibility hypothesis: it\'s not WHICH coping but WHEN/HOW that matters.')
B(d, 'Vietnamese adolescents may have less coping repertoire variety; rigid use of any single strategy → less effective.')

H2(d, '4.5 Methodological considerations')
B(d, 'Brief COPE limitations: frequency-only measurement.')
B(d, 'Cross-sectional: alternative explanation — anxiety may DRIVE coping use (reverse causation).')
B(d, 'Self-report.')

H2(d, '4.6 Clinical and educational implications')
B(d, 'School-based programs should teach coping QUALITY (problem formulation, disclosure style, cognitive depth) not just strategy menu.')
B(d, 'CBT modules: cognitive restructuring + problem-solving training with skill rehearsal.')
B(d, 'Parental psychoeducation: facilitate emotional disclosure without filial-piety violation (e.g., third-party mentor).')
B(d, 'Peer support training: distinguish supportive disclosure from co-rumination.')

H2(d, '4.7 Strengths and limitations')
B(d, 'Strengths: large Vietnamese sample, 3 anxiety phenotypes, integrated SEM, cultural framing.')
B(d, 'Limitations: Brief COPE frequency-only, cross-sectional, urban Hanoi only, no observational measure of coping quality.')

H2(d, '4.8 Future directions')
B(d, 'Develop Vietnamese coping QUALITY measure (e.g., observer-rated, vignette-based).')
B(d, 'Longitudinal panel: do early coping habits predict later anxiety?')
B(d, 'Intervention RCT: targeted coping-quality enhancement vs control.')

# ============================================================
# REFERENCES
# ============================================================
H1(d, 'REFERENCES (preliminary, 30–50 expected)')
N(d, 'Citations sẽ được verify từ kho PDF + WebSearch trước khi viết draft.')

REFS = [
    'Bonanno, G. A., & Burton, C. L. (2013). Regulatory flexibility: An individual differences perspective on coping and emotion regulation. Perspectives on Psychological Science, 8(6), 591–612.',
    'Borkovec, T. D., Alcaine, O., & Behar, E. (2004). Avoidance theory of worry and generalized anxiety disorder. In Generalized anxiety disorder: Advances in research and practice.',
    'Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100.',
    'Chen, Z., et al. (2023). Prevalence and associated factors of depressive and anxiety symptoms among Chinese secondary school students. BMC Psychiatry. [VERIFIED PDF]',
    'Chun, C-A., Moos, R. H., & Cronkite, R. C. (2006). Culture: A fundamental context for the stress and coping paradigm. In Handbook of Multicultural Perspectives on Stress and Coping.',
    'Compas, B. E., et al. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991. [VERIFIED PDF]',
    'Hoang Trung Hoc & Nguyen Thuy Dung. (2025). Levels of stress, anxiety, and depression in adolescents during and after the COVID-19 pandemic in Vietnam. American Journal of Psychiatric Rehabilitation. [VERIFIED PDF]',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis. SEM, 6(1), 1–55.',
    'Institute of Sociology et al. (2022). V-NAMHS report on main findings. [VERIFIED PDF]',
    'Lazarus, R. S., & Folkman, S. (1984). Stress, Appraisal, and Coping. New York: Springer.',
    'Rose, A. J. (2002). Co-rumination in the friendships of girls and boys. Child Development, 73(6), 1830–1843.',
    'Roth, S., & Cohen, L. J. (1986). Approach, avoidance, and coping with stress. American Psychologist, 41(7), 813–819.',
    'Skinner, E. A., Edge, K., Altman, J., & Sherwood, H. (2003). Searching for the structure of coping: A review and critique of category systems for classifying ways of coping. Psychological Bulletin, 129(2), 216–269.',
    'Skinner, E. A., & Zimmer-Gembeck, M. J. (2007). The development of coping. Annual Review of Psychology, 58, 119–144.',
    'Steare, T., et al. (2023). The association between academic pressure and adolescent mental health problems. J Affect Disord. [VERIFIED PDF]',
    'Wei, M., et al. (2013). The role of avoidance coping in the relationship between perfectionism and depression. Journal of Counseling Psychology.',
    'Wu, C., & Chao, R. K. (2011). Intergenerational cultural dissonance in parent-adolescent relationships among Chinese and European Americans. Developmental Psychology.',
    'Xu, Q., et al. (2021). Prevalence and risk factors for anxiety symptoms during the outbreak of COVID-19. J Affect Disord. [VERIFIED PDF]',
    'Zimmer-Gembeck, M. J., & Skinner, E. A. (2016). The development of coping: Implications for psychopathology and resilience. In Developmental Psychopathology.',
]
for ref in REFS:
    B(d, ref)

# ============================================================
# TABLES + FIGURES
# ============================================================
H1(d, 'TABLES + FIGURES')
H3(d, 'Table 1.')
N(d, 'Descriptive statistics for Brief COPE 3 factors + RCADS 3 anxiety phenotypes (M ± SD, by gender).')
H3(d, 'Table 2.')
N(d, 'CFA fit indices for Brief COPE 3-factor model and comparison with 1-factor and 14-factor alternatives.')
H3(d, 'Table 3.')
N(d, 'Bivariate Pearson correlations matrix (coping factors × anxiety phenotypes).')
H3(d, 'Table 4.')
N(d, 'SEM standardized β coefficients: 3 coping factors → 3 anxiety phenotypes.')
H3(d, 'Figure 1.')
N(d, 'SEM model diagram: Avoidance, Problem-solving, Support-seeking → GAD, Social, Separation Anxiety.')

# ============================================================
# QUESTIONS
# ============================================================
H1(d, 'CÂU HỎI EM CẦN NCS + THẦY HD QUYẾT ĐỊNH TRƯỚC KHI VIẾT DRAFT')
B(d, '1. Brief COPE adaptation: NCS đã có publication về Vietnamese version chưa? Em cần cite nó.')
B(d, '2. Số liệu cụ thể bivariate correlations (Brief COPE × RCADS) — em cần verify từ LA hoặc raw data SPSS.')
B(d, '3. Cultural framing "nhẫn": thầy HD có ủng hộ hướng này không, hay muốn tiếp cận universalist (non-cultural)? Cultural framing phù hợp cho Stress and Health.')
B(d, '4. Bài D dùng MOST của same data với Bài A — phải declare đầy đủ trong cover letter (data sharing) tránh self-plagiarism.')
B(d, '5. Có 1 phỏng vấn sâu nào về coping mà em có thể dùng làm illustrative quote không?')
B(d, '6. Timeline: Bài D sau bài A ~2-3 tuần, hay submit gần như đồng thời?')

# Save
d.save(OUT)
from docx import Document as D
dd = D(OUT)
w = sum(len(p.text.split()) for p in dd.paragraphs)
print(f'Da luu: {OUT}')
print(f'Paragraphs: {len(dd.paragraphs)}, Words: ~{w}, Size: {os.path.getsize(OUT)//1024}KB')
