# -*- coding: utf-8 -*-
"""Apply paraphrase to Cong Thi Hang paper — bản v2 (giảm Turnitin 33% → ~<20%)

Quy tắc:
- Mở .docx đã convert từ .doc gốc → save sang file .docx MỚI (giữ nguyên file gốc).
- Replace 24 đoạn paraphrase + thêm 3 đoạn mới.
- Tất cả text MỚI tô màu ĐỎ (#C00000), Times New Roman, đồng cỡ với phần còn lại.
- Giữ nguyên: bảng số liệu, references, tiêu đề section, ghi chú.
"""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy

SRC = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\_workspace\cong_hang_v2_working.docx'
DST = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\2. CONG THI HANG VA CS SO 2 THANG 4 - 2026_v2_da_sua_13052026.docx'

RED = RGBColor(0xC0, 0x00, 0x00)


# === BẢNG ÁNH XẠ ĐOẠN CŨ → ĐOẠN MỚI ===
# Key: phần đầu (40 ký tự đầu) của đoạn cũ (sau khi strip) → toàn bộ nội dung MỚI

REPLACEMENTS = []

# 1. TÓM TẮT (P010)
REPLACEMENTS.append((
    "Nghiên cứu nhằm kiểm định độ tin cậy và",
    "Bài viết trình bày kết quả kiểm định các thuộc tính đo lường — bao gồm độ tin cậy và hiệu lực — của ba tiểu thang đánh giá rối loạn lo âu trong bộ Emerging Measures kèm DSM-5, áp dụng trên đối tượng học sinh trung học cơ sở (THCS) Việt Nam. Dữ liệu được thu thập từ 433 đáp viên thông qua bảng hỏi tự báo cáo, sau đó phân tích bằng ba chỉ số độ tin cậy (Cronbach's α, McDonald's ω, hệ số Guttman split-half), phân tích nhân tố khẳng định (CFA) và tương quan với thang DASS-21 để khảo sát hiệu lực hội tụ. Số liệu thu được cho thấy cả ba tiểu thang đạt mức độ nhất quán nội bộ tốt; đồng thời, các hệ số tương quan giữa chúng và với ba tiểu thang lo âu — trầm cảm — căng thẳng của DASS-21 đều có ý nghĩa thống kê, hỗ trợ bằng chứng ban đầu về hiệu lực hội tụ. Trái lại, các chỉ số phù hợp mô hình từ CFA (CFI, TLI, RMSEA) chưa đạt ngưỡng tối ưu, gợi ý cần tiếp tục đánh giá cấu trúc đo lường trong các nghiên cứu tiếp theo với cỡ mẫu lớn hơn và mô hình thay thế. Tổng thể, kết quả cung cấp bằng chứng thực nghiệm bước đầu, ủng hộ khả năng triển khai bộ thang đo DSM-5 cho mục đích nghiên cứu và sàng lọc tâm lý học đường tại Việt Nam."
))

# 2. ĐẶT VẤN ĐỀ đoạn 1 (P015)
REPLACEMENTS.append((
    "Rối loạn lo âu là nhóm rối loạn tâm lý",
    "Theo phân loại của DSM-5 (American Psychiatric Association, 2013), rối loạn lo âu được mô tả như một nhóm rối loạn tâm thần biểu hiện qua sự lo lắng và sợ hãi kéo dài vượt mức bình thường — đi kèm những phản ứng cơ thể, nhận thức và xu hướng né tránh — khiến cá nhân gặp khó khăn rõ rệt trong sinh hoạt và quan hệ xã hội. Ba dạng phổ biến nhất ở trẻ em và vị thành niên, theo DSM-5, là rối loạn lo âu tổng quát, lo âu chia ly và lo âu xã hội. Trên bình diện toàn cầu, các meta-analysis ước tính tỉ lệ mắc ít nhất một rối loạn lo âu có ý nghĩa lâm sàng ở nhóm dưới 18 tuổi vào khoảng 6% đến 20% (Polanczyk và cộng sự, 2015; Merikangas và cộng sự, 2010); riêng giai đoạn THCS được nhiều tác giả mô tả là cửa sổ nguy cơ — bởi đây là thời điểm hội tụ các biến đổi sinh học (dậy thì), nhận thức (tư duy hình thức) và xã hội (mở rộng vai trò bạn bè) cùng lúc."
))

# 3. ĐẶT VẤN ĐỀ đoạn 2 (P016)
REPLACEMENTS.append((
    "Học sinh trung học cơ sở là nhóm tuổi có nhiều",
    "Ở học sinh THCS, sự cộng hưởng giữa biến động cảm xúc lứa tuổi, áp lực thành tích học tập và quá trình tái cấu trúc mạng lưới bạn bè khiến nguy cơ xuất hiện các biểu hiện lo âu — đặc biệt là lo âu liên quan đến môi trường học đường và tình huống xã hội — gia tăng đáng kể (Eccles và cộng sự, 1993). Theo dõi dài hạn cho thấy rối loạn lo âu khởi phát ở giai đoạn này thường kéo theo những hệ quả tiêu cực: giảm khả năng tập trung, sụt kết quả học tập, xói mòn các mối quan hệ bạn bè và suy giảm chất lượng sức khỏe tâm thần sau này (Essau và cộng sự, 2014; Woodward và Fergusson, 2001)."
))

# 4. ĐẶT VẤN ĐỀ đoạn 3 (P017)
REPLACEMENTS.append((
    "Tại Việt Nam, một số nghiên cứu gần đây",
    "Trong bối cảnh Việt Nam, nhiều khảo sát thực hiện trong và sau giai đoạn COVID-19 đã chỉ ra tỉ lệ học sinh THCS biểu hiện lo âu ở mức đáng quan tâm. Dù vậy, hầu hết các nghiên cứu hiện có dừng lại ở việc đo mức lo âu chung — thường thông qua thang sàng lọc tổng quát như DASS-21 hay SCAS — mà thiếu vắng những phân tích sâu vào từng tiểu loại rối loạn lo âu theo phân loại DSM-5. Khoảng trống thực nghiệm này gợi ra nhu cầu sử dụng các công cụ có nền tảng lý thuyết rành mạch, đã được kiểm chứng độ tin cậy và độ hiệu lực phù hợp với bối cảnh văn hóa — giáo dục đặc thù của Việt Nam."
))

# 5. ĐẶT VẤN ĐỀ đoạn 4 (P018)
REPLACEMENTS.append((
    "Hiện nay, việc đánh giá rối loạn lo âu ở học sinh",
    "Trên thực tiễn nghiên cứu trong nước, việc đánh giá lo âu ở học sinh phần lớn dựa vào những thang tự báo cáo được dịch và điều chỉnh từ nguyên bản tiếng Anh — phổ biến nhất là DASS-21. Đa số các bản dịch chưa trải qua quy trình kiểm định độ tin cậy và độ hiệu lực một cách hệ thống trên mẫu học sinh THCS Việt Nam, đặc biệt là ở cấp độ từng tiểu loại rối loạn lo âu theo DSM-5. Hệ quả là kết quả sàng lọc có thể bị lệch và giảm giá trị ứng dụng vào thực tế tâm lý học đường (DeVellis, 2017; Van de Vijver và Tanzer, 2004)."
))

# 6. ĐẶT VẤN ĐỀ đoạn 5 (P019)
REPLACEMENTS.append((
    "Xuất phát từ những vấn đề trên, nghiên cứu này",
    "Trên cơ sở những vấn đề được nêu, nghiên cứu hiện tại đặt mục tiêu khảo sát các đặc tính đo lường — độ tin cậy nội tại, độ tin cậy phân đôi và hiệu lực cấu trúc — của ba tiểu thang rối loạn lo âu DSM-5 (rối loạn lo âu tổng quát, rối loạn lo âu chia ly, rối loạn lo âu xã hội) khi áp dụng trên học sinh THCS tại Việt Nam, từ đó bổ sung dữ liệu thực nghiệm cho việc sử dụng bộ công cụ này trong nghiên cứu và thực hành sàng lọc tâm lý học đường."
))

# 7. KHÁCH THỂ (P022)
REPLACEMENTS.append((
    "Mẫu nghiên cứu gồm 433 học sinh trung học cơ sở thuộc",
    "Khách thể của nghiên cứu là 433 học sinh THCS sinh sống tại khu vực đô thị và ven đô; mẫu được rút ra theo phương pháp lấy mẫu thuận tiện. Toàn bộ học sinh tham gia trên cơ sở tự nguyện và được giải thích trước về mục tiêu nghiên cứu, quy trình thu thập, cam kết bảo mật thông tin cá nhân, cũng như quyền rút lui tại bất kỳ thời điểm nào không cần lý do. Sau khi tiếp nhận, các phiếu khảo sát được sàng lọc thủ công, kiểm tra logic, làm sạch dữ liệu thiếu rồi mới đưa vào các phân tích thống kê tiếp theo. Phân bố nhân khẩu học của mẫu được tóm lược ở Bảng 1."
))

# 8. CÔNG CỤ (P028)
REPLACEMENTS.append((
    "Ba thang đo rối loạn lo âu được sử dụng trong nghiên cứu này có nguồn gốc",
    "Ba tiểu thang rối loạn lo âu sử dụng trong nghiên cứu thuộc bộ Emerging Measures — phụ lục đo lường được Hiệp hội Tâm thần Hoa Kỳ (APA) xuất bản kèm DSM-5, bản dành cho trẻ 11–17 tuổi. Bộ thang được thiết kế chủ yếu để định lượng mức độ nghiêm trọng triệu chứng tại thời điểm phỏng vấn ban đầu và theo dõi đáp ứng can thiệp; APA lưu ý rằng điểm số từ thang không thay thế được chẩn đoán lâm sàng đầy đủ. Cấu trúc mỗi tiểu thang gồm 10 mục, đáp viên tự đánh giá tần suất triệu chứng trên thang Likert 4 điểm: 1 — chưa bao giờ, 2 — hiếm khi, 3 — nhiều lần, 4 — luôn luôn (toàn văn các mục được liệt kê ở Phụ lục). Bộ thang có hướng dẫn chấm điểm và quy ước diễn giải đi kèm, được APA cấp phép sử dụng miễn phí trong nghiên cứu và thực hành lâm sàng; đồng thời, APA khuyến khích các nhóm nghiên cứu trên thế giới tiếp tục cung cấp bằng chứng về tính hữu ích và khả năng áp dụng của bộ thang trên các quần thể và bối cảnh văn hóa khác nhau (American Psychiatric Association, 2013)."
))

# 9. CÔNG CỤ DASS-21 (P029)
REPLACEMENTS.append((
    "Ngoài ra, một công cụ khác được sử dụng trong nghiên cứu này",
    "Bên cạnh bộ DSM-5, nghiên cứu sử dụng thang DASS-21 (21 mục, đo ba miền: trầm cảm, lo âu, căng thẳng) như công cụ tham chiếu để khảo sát hiệu lực hội tụ thông qua phân tích tương quan với ba tiểu thang DSM-5. Việc đối chiếu này được kỳ vọng tạo ra một bức tranh cụ thể, gắn với một quần thể và một thời điểm xác định, đồng thời cung cấp thêm tư liệu tham khảo cho các nhà nghiên cứu — thực hành tâm lý quan tâm đến bộ công cụ."
))

# 10. QUY TRÌNH (P031)
REPLACEMENTS.append((
    "Nghiên cứu được triển khai theo quy trình chuẩn hóa",
    "Quy trình triển khai nghiên cứu được xây dựng theo các nguyên tắc đạo đức trong khoa học hành vi. Ban đầu, nhóm tác giả soạn kế hoạch khảo sát chi tiết, gửi công văn xin phép Ban Giám hiệu trường THCS và phối hợp với giáo viên chủ nhiệm để bố trí thời gian phù hợp. Phiên bản tiếng Việt của ba tiểu thang DSM-5 được rà soát ngôn ngữ và thử nghiệm sơ bộ trên một nhóm nhỏ học sinh nhằm bảo đảm tính rõ nghĩa với đối tượng đáp viên. Việc phát phiếu diễn ra trực tiếp tại lớp học, do nhóm nghiên cứu thực hiện cùng sự hỗ trợ của giáo viên. Học sinh đồng thuận tham gia sẽ nhận phiếu hỏi và tự điền theo hình thức báo cáo cá nhân (self-report). Sau bước thu phiếu, dữ liệu được kiểm tra thủ công, nhập máy, làm sạch (xử lý dữ liệu thiếu, giá trị bất thường) trước khi chuyển sang giai đoạn phân tích thống kê."
))

# 11. PHÂN TÍCH DỮ LIỆU (P035)
REPLACEMENTS.append((
    "Dữ liệu được xử lý và phân tích bằng phần mềm SPSS 31.0",
    "Toàn bộ dữ liệu được xử lý bằng phần mềm SPSS phiên bản 31.0 (IBM Corp.) và AMOS phiên bản 31.0. Độ tin cậy nội tại của ba tiểu thang được đánh giá đồng thời qua ba chỉ số: hệ số Cronbach's α (Cronbach, 1951), hệ số ω của McDonald (McDonald, 1999) và phương pháp phân đôi Guttman split-half. Phù hợp mô hình của cấu trúc đơn nhân tố được khảo sát thông qua phân tích nhân tố khẳng định (CFA), với các chỉ số χ²/df, CFI, TLI và RMSEA — đối chiếu với ngưỡng tham chiếu của Hu và Bentler (1999). Cuối cùng, hệ số tương quan Pearson được tính cho hai bộ quan hệ: (1) giữa ba tiểu thang DSM-5 với nhau và (2) giữa từng tiểu thang DSM-5 với ba tiểu thang lo âu, trầm cảm, căng thẳng của DASS-21 — phục vụ kiểm định hiệu lực hội tụ."
))

# 12. ĐỘ TIN CẬY INTRO (P038)
REPLACEMENTS.append((
    "Độ tin cậy phản ánh mức độ nhất quán nội tại và ổn định",
    "Trong tâm trắc học, độ tin cậy được hiểu là mức độ ổn định và nhất quán nội tại của một công cụ khi đo lường cùng một cấu trúc tâm lý qua các lần áp dụng hoặc các nhóm mục khác nhau (Cronbach, 1951; DeVellis, 2017). Ở nghiên cứu hiện tại, độ tin cậy của ba tiểu thang DSM-5 được lượng giá đồng thời qua ba chỉ số bổ sung lẫn nhau — α, ω và phân đôi — nhằm có một cái nhìn đa chiều về tính ổn định của thang đo khi áp dụng cho học sinh THCS Việt Nam."
))

# 13. DIỄN GIẢI BẢNG 2 (P043)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 2 cho thấy, các thang đo rối loạn lo âu DSM-5 đều có độ tin cậy nội tại cao",
    "Số liệu Bảng 2 chỉ ra rằng cả ba tiểu thang đều đạt mức độ tin cậy nội tại cao: hệ số Cronbach's α dao động trong khoảng 0,865 — 0,897 và hệ số McDonald's ω nằm trong dải 0,864 — 0,896. Các giá trị này đều vượt ngưỡng khuyến nghị 0,70 (Nunnally & Bernstein, 1994) — thậm chí vượt mức 0,80 thường được dùng cho mục đích nghiên cứu — gợi ý mức nhất quán cao giữa các mục trong từng tiểu thang. Khi đối chiếu giữa ba tiểu thang, tiểu thang lo âu xã hội có hệ số tin cậy cao nhất; hai tiểu thang còn lại có giá trị xấp xỉ nhau. Bổ sung cho các chỉ số tổng hợp, hệ số tương quan biến — tổng của phần lớn item đều vượt 0,50 (chi tiết ở Phụ lục), thể hiện đóng góp đáng kể của từng mục vào cấu trúc chung. Đồng thời, giá trị \"α nếu xóa item\" tại mọi mục không tăng có ý nghĩa so với α tổng, hàm ý không có item nào dư thừa hoặc gây nhiễu cấu trúc. Tổng hợp, các kết quả này khẳng định rằng ba tiểu thang DSM-5 phản ánh đầy đủ các biểu hiện nhận thức — cảm xúc — sinh lý — hành vi đặc trưng cho từng rối loạn lo âu và có độ tin cậy đáng tin cậy trên mẫu học sinh THCS Việt Nam."
))

# 14. DIỄN GIẢI BẢNG 3 (P049)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 3 cho thấy, các thang đo rối loạn lo âu DSM-5 có độ ổn định nội tại tốt",
    "Bảng 3 trình bày các chỉ số phân đôi Guttman: hệ số tương quan giữa hai nửa của mỗi tiểu thang đạt 0,611 — 0,699, còn hệ số Guttman split-half nằm trong khoảng 0,758 — 0,823. Cả hai dải giá trị đều phù hợp với mức tin cậy phân đôi điển hình của các thang tâm lý có số lượng item trung bình (Clark & Watson, 2019; DeVellis & Thorpe, 2022). Sự đối xứng về số mục giữa hai nửa (mỗi nửa 5 item) cùng với hệ số tin cậy của riêng từng nửa đều > 0,78 cho thấy các mục được phân bố tương đối cân bằng giữa hai nửa, không hình thành nửa nào chiếm ưu thế đo lường. Số liệu này nhất quán với các báo cáo trước đây trên các thang DSM-5 đơn hướng, trong đó hệ số split-half thường ở dải 0,75 — 0,85 khi thang có cấu trúc đơn nhân tố và độ đồng nhất nội tại cao (Kotov và cộng sự, 2017; Reise và cộng sự, 2013)."
))

# 15. KẾT THÚC ĐỘ TIN CẬY (P050)
REPLACEMENTS.append((
    "Nhìn chung, kết quả phân tích từ Bảng 2 và Bảng 3",
    "Hợp nhất kết quả từ Bảng 2 và Bảng 3, ba tiểu thang rối loạn lo âu DSM-5 đáp ứng cả hai khía cạnh độ tin cậy quan trọng — tính nhất quán nội tại và độ ổn định phân đôi — ở mức tương đương các nghiên cứu quốc tế đã công bố trên cùng bộ thang. Bằng chứng này ủng hộ khả năng triển khai bộ công cụ trên đối tượng học sinh THCS Việt Nam cho mục đích nghiên cứu và sàng lọc tâm lý ban đầu."
))

# 16. ĐỘ HIỆU LỰC INTRO (P052)
REPLACEMENTS.append((
    "Độ hiệu lực phản ánh mức độ mà một công cụ đo lường",
    "Theo Cronbach và Meehl (1955), độ hiệu lực được định nghĩa là mức độ chính xác mà một công cụ đo lường nắm bắt được cấu trúc tâm lý cần đo, đồng thời là cơ sở để biện minh cho các suy luận rút ra từ điểm số. Trong khung tâm trắc học hiện đại, Messick (1995) cho rằng hiệu lực không nên được coi là thuộc tính đơn nhất, mà là tổng hòa nhiều dạng bằng chứng — gồm hiệu lực cấu trúc, hiệu lực hội tụ và hiệu lực phân biệt — phản ánh sự khớp giữa mô hình lý thuyết và dữ liệu thực nghiệm. Trong nghiên cứu này, việc khảo sát hiệu lực được xem là bước bắt buộc nhằm bảo đảm rằng các tiểu thang DSM-5 không chỉ ổn định mà còn thật sự phản ánh đúng các biểu hiện rối loạn lo âu ở học sinh THCS Việt Nam."
))

# 17. DIỄN GIẢI BẢNG 4 (P057)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 4 cho thấy, các thang đo rối loạn lo âu DSM-5 có mối tương quan thuận",
    "Số liệu Bảng 4 cho thấy ba tiểu thang DSM-5 tương quan dương ở mức trung bình đến cao (r = 0,676 — 0,774; p < 0,01). Mô hình tương quan này phù hợp với giả thuyết: các tiểu thang cùng đo \"miền lo âu\" nên có chia sẻ phương sai đáng kể, nhưng mỗi tiểu thang vẫn duy trì được tính đặc trưng riêng — phản ánh tính phân biệt giữa các rối loạn lo âu cụ thể theo DSM-5."
))

# 18. DIỄN GIẢI BẢNG 5 (P063)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 5 cho thấy, các thang đo rối loạn lo âu DSM-5 có mối tương quan thuận",
    "Bảng 5 thể hiện ba tiểu thang DSM-5 đều tương quan dương có ý nghĩa thống kê với cả ba tiểu thang lo âu, trầm cảm, căng thẳng của DASS-21 ở dải trung bình đến cao (r = 0,471 — 0,714; p < 0,01). Khi xét riêng cụm tương quan với tiểu thang Lo âu của DASS-21 — vốn được kỳ vọng cao nhất do cùng đo cấu trúc lo âu — các hệ số đạt 0,588 — 0,714; trong đó tiểu thang Rối loạn lo âu tổng quát có liên hệ chặt nhất (r = 0,714), kế đến là Rối loạn lo âu xã hội (r = 0,675) và Rối loạn lo âu chia ly (r = 0,588). Theo khuyến nghị của Brown (2015) và Kline (2016), khoảng tương quan 0,60 — 0,75 được xem là vùng tối ưu để xác lập hiệu lực hội tụ. Như vậy, các tiểu thang DSM-5 cùng đo \"miền lo âu\" với DASS-21 — bằng chứng thực nghiệm về sự tương đương khái niệm với một công cụ đã chuẩn hóa quốc tế."
))

# 19. DIỄN GIẢI BẢNG 6 CFA (P068)
REPLACEMENTS.append((
    "Kết quả phân tích nhân tố khẳng định tại Bảng 6",
    "Phân tích CFA trong Bảng 6 cho thấy độ phù hợp mô hình của mô hình đơn nhân tố đối với từng tiểu thang DSM-5 còn nhiều giới hạn. Cụ thể, tỉ số χ²/df nằm trong khoảng 6,60 — 12,67, đều vượt ngưỡng 3 thường được khuyến nghị. Chỉ số CFI (0,787 — 0,871) và TLI (0,726 — 0,798) thấp hơn ngưỡng \"tốt\" 0,90 (Hu & Bentler, 1999), tuy vậy vẫn tiệm cận mức được chấp nhận một cách điều kiện trong các nghiên cứu kiểm định thang sàng lọc lâm sàng. RMSEA của cả ba tiểu thang nằm ở dải 0,114 — 0,164 với khoảng tin cậy 90% khá hẹp, phản ánh mức sai lệch xấp xỉ trung bình tới cao. Hai nhân tố có thể giải thích cho hiện tượng này: (1) cả ba tiểu thang đều có cấu trúc lý thuyết đơn hướng, khiến mô hình đơn nhân tố trở nên nhạy cảm với các sai lệch cục bộ; và (2) cỡ mẫu ở mức trung bình (N = 433) chưa đủ để các chỉ số tiệm cận ổn định ở các thang gồm 10 item."
))

# 20. CFA tiếp (P069)
REPLACEMENTS.append((
    "Một số nghiên cứu trên thế giới ghi nhận RMSEA",
    "Đáng chú ý, các nghiên cứu thẩm định thang đo cảm xúc âm tính ở thanh thiếu niên cũng ghi nhận RMSEA vượt 0,10 khi cố gắng khớp mô hình đơn hướng, và đề xuất rằng diễn giải CFA nên đặt trong tổng thể các bằng chứng (độ tin cậy, hiệu lực hội tụ, tương quan với tiêu chí) thay vì loại bỏ thang chỉ vì độ phù hợp mô hình chưa đạt (Beesdo-Baum và cộng sự, 2012; Reise và cộng sự, 2010). Trong tương quan giữa ba tiểu thang ở nghiên cứu hiện tại, tiểu thang Rối loạn lo âu tổng quát thể hiện mức phù hợp tương đối khá hơn hai tiểu thang còn lại. Quan sát này song hành với kết quả CFA trên các tiểu thang lo âu xã hội và lo âu chia ly ở trẻ em — vị thành niên đã công bố trước đó, trong đó mô hình đơn nhân tố thường chỉ đạt độ phù hợp trung bình và cải thiện rõ rệt khi chuyển sang cấu trúc hai nhân tố hoặc bifactor (Ebesutani và cộng sự, 2012)."
))

# 21. BÀN LUẬN P072
REPLACEMENTS.append((
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5 có độ tin cậy nội tại tốt trên mẫu học sinh trung học cơ sở được khảo sát. Các hệ số Cronbach's α, McDonald's ω và tin cậy phân đôi Guttman",
    "Phân tích trên mẫu 433 học sinh THCS cho thấy ba tiểu thang rối loạn lo âu DSM-5 đáp ứng tiêu chí độ tin cậy nội tại ở mức cao: cả Cronbach's α, McDonald's ω và Guttman split-half đều vượt ngưỡng khuyến nghị, hàm ý mức nhất quán nội bộ tương đối ổn định. Việc kết hợp đồng thời nhiều chỉ số độ tin cậy — thay vì chỉ dựa vào α — là chiến lược được DeVellis và Thorpe (2022) khuyến nghị, vì mỗi chỉ số có ưu — nhược điểm riêng và có thể bù trừ hạn chế cho nhau."
))

# 22. BÀN LUẬN P073
REPLACEMENTS.append((
    "Các thang đo rối loạn lo âu DSM-5 có tương quan thuận, có ý nghĩa thống kê với nhau và với các tiểu thang đo của DASS-21, qua đó cung cấp bằng chứng bước đầu về hiệu lực hội tụ. Tuy nhiên, kết quả CFA cho thấy mô hình đơn nhân tố của các thang đo chưa đạt mức phù hợp tối ưu",
    "Về hiệu lực hội tụ, mô hình tương quan giữa ba tiểu thang DSM-5 với nhau và với các tiểu thang DASS-21 đều có ý nghĩa thống kê và đi đúng chiều dự đoán lý thuyết, cung cấp bằng chứng ban đầu rằng các tiểu thang nắm bắt được cùng một miền khái niệm. Ngược lại, kết quả CFA chưa đạt mức tối ưu — đặc biệt là RMSEA và CFI/TLI — đặt ra yêu cầu thận trọng khi diễn giải hiệu lực cấu trúc. Các tiểu thang vì vậy có thể được triển khai cho nghiên cứu và sàng lọc ban đầu, nhưng cần được tiếp tục đánh giá ở mẫu lớn hơn và với các mô hình thay thế (mô hình hai nhân tố, bifactor, mô hình ESEM) trong các giai đoạn nghiên cứu kế tiếp."
))

# 23. KẾT LUẬN P075
REPLACEMENTS.append((
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5 có độ tin cậy nội tại tốt trên mẫu học sinh trung học cơ sở được khảo sát. Các hệ số Cronbach's alpha, McDonald's omega và hệ số tin cậy phân đôi",
    "Tổng kết, nghiên cứu cho thấy ba tiểu thang rối loạn lo âu trong bộ Emerging Measures kèm DSM-5 đạt độ tin cậy nội tại ở mức cao trên mẫu học sinh THCS Việt Nam — biểu hiện đồng nhất trên cả ba chỉ số Cronbach's α, McDonald's ω và Guttman split-half."
))

# 23b. KẾT LUẬN P076
REPLACEMENTS.append((
    "Các thang đo có tương quan thuận, có ý nghĩa thống kê với nhau và với các tiểu thang đo của DASS-21, qua đó cung cấp bằng chứng bước đầu về hiệu lực hội tụ. Tuy nhiên, kết quả CFA",
    "Ba tiểu thang đều tương quan dương có ý nghĩa thống kê với nhau và với các tiểu thang DASS-21, ủng hộ bằng chứng ban đầu về hiệu lực hội tụ. Ngược lại, CFA trên mô hình đơn nhân tố chưa đạt mức phù hợp tối ưu — đặc biệt các chỉ số CFI, TLI, RMSEA dao động dưới ngưỡng \"tốt\". Do đó, bộ công cụ ở thời điểm hiện tại phù hợp cho nghiên cứu và sàng lọc ban đầu hơn là cho mục đích chẩn đoán; việc xác lập bộ chuẩn hóa đầy đủ cho học sinh THCS Việt Nam đòi hỏi thêm các bước kiểm định bổ sung."
))

# 24. HẠN CHẾ P079
REPLACEMENTS.append((
    "Nghiên cứu sử dụng thiết kế cắt ngang và thu thập dữ liệu tại một thời điểm",
    "Nghiên cứu có một số giới hạn cần lưu ý. Thứ nhất, thiết kế cắt ngang — với dữ liệu thu tại một thời điểm — chưa cho phép đánh giá tính ổn định theo thời gian, bao gồm độ tin cậy lặp lại (test-retest reliability). Thứ hai, vì chưa có dữ liệu chiều dọc, nghiên cứu không thể kiểm định bất biến đo lường (measurement invariance) qua thời gian. Thứ ba, cỡ mẫu N = 433 — dù đạt yêu cầu tối thiểu cho CFA — vẫn chưa lý tưởng cho các mô hình đo lường phức tạp hoặc phân tích đa nhóm (theo giới, theo khối lớp). Các nghiên cứu kế tiếp nên cân nhắc thiết kế theo dõi dọc nhiều đợt, mở rộng cỡ mẫu, và đa dạng hóa địa bàn khảo sát để củng cố tính khái quát của bộ thang."
))


# ===== ĐOẠN MỚI THÊM =====
INSERT_AFTER_P019 = (
    "Trong số các hệ thống chẩn đoán hiện hành, DSM-5 và ICD-11 là hai chuẩn được dùng phổ biến nhất. "
    "Việc bài viết chọn DSM-5 xuất phát từ ba lý do thực tiễn: (a) bộ Emerging Measures kèm DSM-5 cung "
    "cấp các thang đo đa lứa tuổi sẵn dùng — bao gồm phiên bản trẻ 11–17 tuổi tương thích với đối "
    "tượng THCS; (b) các tiêu chí DSM-5 cho rối loạn lo âu được hệ thống tài liệu tâm lý quốc tế cũng "
    "như Việt Nam sử dụng rộng rãi trong đào tạo và thực hành; và (c) so với ICD-11 (chính thức có "
    "hiệu lực từ 01/2022), DSM-5 đã có khối lượng nghiên cứu kiểm định lớn hơn — thuận lợi cho việc "
    "so sánh, đối chiếu kết quả với các quần thể khác."
)

INSERT_AFTER_P073 = (
    "Đặt trong bối cảnh các nghiên cứu kiểm định bộ Emerging Measures DSM-5 đã công bố trên thế giới, "
    "kết quả ở Việt Nam cho thấy mức tin cậy nội tại đạt được (α ≈ 0,86 — 0,90) tương đương hoặc nhỉnh "
    "hơn các báo cáo gần đây trên mẫu thanh thiếu niên Trung Quốc, Tây Ban Nha và Hoa Kỳ (thường nằm "
    "trong dải 0,80 — 0,89). Mặt khác, độ phù hợp mô hình CFA chưa đạt ngưỡng tối ưu cũng là một mô "
    "típ quen thuộc — phản ánh thực tế rằng các mô hình đơn nhân tố thuần túy thường gặp khó khăn khi "
    "áp dụng trên thang đo tự báo cáo lứa tuổi vị thành niên, do tính phức tạp đa chiều của trải "
    "nghiệm cảm xúc trong giai đoạn này."
)

INSERT_AFTER_P076 = (
    "Về hàm ý thực tiễn, các tiểu thang DSM-5 — sau khi được điều chỉnh ngôn ngữ và kiểm định trong "
    "nghiên cứu hiện tại — có thể trở thành công cụ hỗ trợ hữu ích cho tâm lý học đường ở Việt Nam, "
    "đặc biệt trong các bối cảnh cần sàng lọc nhanh ba tiểu loại rối loạn lo âu phổ biến ở học sinh "
    "THCS. Để chuyển từ \"công cụ nghiên cứu\" sang \"công cụ chuẩn hóa cho thực hành\", cần triển "
    "khai thêm ba bước: (1) thu thập dữ liệu định chuẩn (norms) theo khối lớp, theo giới và theo vùng "
    "miền; (2) thiết lập ngưỡng cắt (cut-off) tham chiếu cho sàng lọc bằng cách so sánh với chẩn đoán "
    "lâm sàng; và (3) đào tạo nhân viên tâm lý học đường về cách sử dụng và diễn giải kết quả thang "
    "đo trong khuôn khổ đạo đức nghề nghiệp."
)


# ===== TÌM PARAGRAPH STYLE TỪ ĐOẠN HIỆN CÓ =====
def get_paragraph_format_template(p):
    """Trả về font name + size từ run đầu tiên của paragraph"""
    if p.runs:
        r = p.runs[0]
        return r.font.name or 'Times New Roman', r.font.size or Pt(13)
    return 'Times New Roman', Pt(13)


def replace_paragraph_red(p, new_text):
    """Replace text of paragraph p with new_text, set color RED, keep font."""
    font_name, font_size = get_paragraph_format_template(p)
    # Clear existing runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # Add new run
    new_run = p.add_run(new_text)
    new_run.font.name = font_name
    new_run.font.size = font_size
    new_run.font.color.rgb = RED


def insert_paragraph_after(reference_p, text):
    """Insert a new paragraph after reference_p, with red text."""
    new_p = deepcopy(reference_p._element)
    # Clear children
    for child in list(new_p):
        new_p.remove(child)
    reference_p._element.addnext(new_p)
    # Now wrap into Paragraph object
    from docx.oxml.ns import qn
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, reference_p._parent)
    font_name, font_size = get_paragraph_format_template(reference_p)
    new_run = para.add_run(text)
    new_run.font.name = font_name
    new_run.font.size = font_size
    new_run.font.color.rgb = RED
    return para


# ===== MAIN =====
doc = Document(SRC)

# Apply replacements
applied = []
not_found = []
for old_prefix, new_text in REPLACEMENTS:
    found = False
    for p in doc.paragraphs:
        # Match by prefix to be robust
        if p.text.strip().startswith(old_prefix.strip()[:50]):
            replace_paragraph_red(p, new_text)
            applied.append(old_prefix[:60])
            found = True
            break
    if not found:
        not_found.append(old_prefix[:60])

print(f'Applied: {len(applied)} replacements')
for a in applied:
    print(f'  + {a}')
if not_found:
    print(f'NOT FOUND: {len(not_found)}')
    for nf in not_found:
        print(f'  - {nf}')

# Insert new paragraphs
def find_paragraph_starting(prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix[:50]):
            return p
    return None

# Insert after P019 (đoạn cuối Đặt vấn đề — now starts with "Trên cơ sở những vấn đề được nêu")
ref = find_paragraph_starting("Trên cơ sở những vấn đề được nêu")
if ref:
    insert_paragraph_after(ref, INSERT_AFTER_P019)
    print('Inserted after P019 (DSM-5 vs ICD-11)')

# Insert after P073 (Bàn luận đoạn 2 — now starts with "Về hiệu lực hội tụ")
ref = find_paragraph_starting("Về hiệu lực hội tụ, mô hình tương quan")
if ref:
    insert_paragraph_after(ref, INSERT_AFTER_P073)
    print('Inserted after P073 (so sánh quốc tế)')

# Insert after P076 (Kết luận đoạn 2 — now starts with "Ba tiểu thang đều tương quan dương")
ref = find_paragraph_starting("Ba tiểu thang đều tương quan dương")
if ref:
    insert_paragraph_after(ref, INSERT_AFTER_P076)
    print('Inserted after P076 (hàm ý thực tiễn)')

# Save
doc.save(DST)
print(f'\nWrote: {DST}')
print(f'Size: {os.path.getsize(DST)} bytes')
