# -*- coding: utf-8 -*-
"""Sinh doc Quy dinh trinh bay LA TS Viet Nam cho thay.
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '10_QuyDinh_TrinhBay_LuanAnTS_27052026.docx')

BLACK = RGBColor(0, 0, 0)
GRAY = RGBColor(80, 80, 80)
BLUE = RGBColor(0, 51, 102)


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in d.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.0)
    cp = d.core_properties
    cp.author = ''; cp.title = ''; cp.subject = ''
    cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
    return d


def H(d, text, level=1):
    h = d.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = BLACK
    return h


def P(d, text, bold=False, italic=False, color=None, indent=True, size=13):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def add_table(d, headers, rows, widths=None):
    n = len(headers)
    t = d.add_table(rows=1, cols=n)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
        if widths:
            hdr[i].width = widths[i]
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            r = p.add_run(cell_data)
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)
            if widths:
                row[i].width = widths[i]
    return t


d = doc_init()

# Title
H(d, 'QUY ĐỊNH TRÌNH BÀY LUẬN ÁN TIẾN SĨ TẠI VIỆT NAM', 1)
P(d, '(Tổng hợp theo Thông tư 18/2021/TT-BGDĐT của Bộ Giáo dục và Đào tạo + quy định phổ biến của các cơ sở đào tạo)', italic=True, indent=False, size=11, color=GRAY)
P(d, '27/05/2026 — NCS Công Thị Hằng', italic=True, indent=False, size=11, color=GRAY)
d.add_paragraph()

P(d, 'Kính gửi thầy,', indent=False)
P(d, 'Em xin tổng hợp quy chuẩn trình bày luận án tiến sĩ phổ biến nhất ở Việt Nam, để thầy đối chiếu khi rà soát bản nháp của em. Mỗi cơ sở đào tạo có thể có điều chỉnh nhỏ, nhưng các thông số dưới đây là chuẩn áp dụng rộng rãi cho khối Khoa học Xã hội — Giáo dục — Tâm lý học.')
d.add_paragraph()

# 1. Dinh dang co ban
H(d, '1. Định dạng cơ bản', 2)
add_table(d,
    ['Mục', 'Quy chuẩn'],
    [
        ('Khổ giấy', 'A4 (210 × 297 mm), in một mặt'),
        ('Phông chữ', 'Times New Roman (Unicode), bộ mã TCVN 6909:2001'),
        ('Cỡ chữ body', '13 pt (một số trường dùng 14 pt)'),
        ('Giãn cách dòng', '1,5 lines (Multiple 1.5)'),
        ('Giãn cách đoạn', 'Before: 0 pt — After: 6 pt (hoặc 0/0 tùy trường)'),
        ('Thụt đầu dòng', '1,0–1,27 cm (First line indent)'),
        ('Căn lề văn bản', 'Justify — căn đều hai bên'),
    ],
    widths=[Cm(5.0), Cm(11.0)])
d.add_paragraph()

# 2. Le trang
H(d, '2. Lề trang (margins)', 2)
add_table(d,
    ['Lề', 'Khoảng cách phổ biến'],
    [
        ('Lề trên (top)', '2,0–2,5 cm'),
        ('Lề dưới (bottom)', '2,0–2,5 cm'),
        ('Lề trái (left) — bên đóng gáy', '3,0–3,5 cm'),
        ('Lề phải (right)', '2,0 cm'),
    ],
    widths=[Cm(7.0), Cm(9.0)])
P(d, 'Quy chuẩn phổ biến nhất: Trên 2,5 / Dưới 2,5 / Trái 3,5 / Phải 2,0 cm.', italic=True, indent=False, size=11)
d.add_paragraph()

# 3. Tieu de cac cap
H(d, '3. Tiêu đề các cấp', 2)
add_table(d,
    ['Cấp', 'Cỡ chữ', 'Kiểu', 'Căn'],
    [
        ('Tên chương (CHƯƠNG 1...)', '16–18 pt', 'IN HOA, đậm', 'Giữa'),
        ('Mục cấp 2 (1.1.)', '14 pt', 'Đậm', 'Trái'),
        ('Mục cấp 3 (1.1.1.)', '13 pt', 'Đậm', 'Trái'),
        ('Mục cấp 4 (1.1.1.1.)', '13 pt', 'Nghiêng đậm', 'Trái'),
    ],
    widths=[Cm(5.5), Cm(3.0), Cm(4.5), Cm(3.0)])
P(d, 'Lưu ý: không nên đánh số sâu hơn 4 cấp để giữ tính súc tích.', italic=True, indent=False, size=11)
d.add_paragraph()

# 4. Bang va hinh
H(d, '4. Bảng và hình', 2)
P(d, 'a) Tiêu đề bảng:', bold=True, indent=False)
P(d, 'Đặt PHÍA TRÊN bảng, in đậm hoặc nghiêng. Định dạng: "Bảng X.Y. Tên bảng" — trong đó X là số chương, Y là thứ tự bảng trong chương.')
P(d, 'b) Tiêu đề hình:', bold=True, indent=False)
P(d, 'Đặt PHÍA DƯỚI hình. Định dạng: "Hình X.Y. Tên hình".')
P(d, 'c) Nguồn:', bold=True, indent=False)
P(d, 'Đặt dưới bảng/hình, in nghiêng, cỡ 11–12 pt. Ví dụ: Nguồn: Tác giả tổng hợp, 2025.')
P(d, 'd) Nội dung bên trong bảng:', bold=True, indent=False)
P(d, 'Có thể dùng cỡ 11–12 pt, giãn dòng 1,0 hoặc 1,15 (compact) để tiết kiệm trang in.')
d.add_paragraph()

# 5. Danh so trang
H(d, '5. Đánh số trang', 2)
P(d, 'Phần đầu (lời cam đoan, lời cảm ơn, mục lục, danh mục bảng/hình, danh mục từ viết tắt):', bold=True, indent=False)
P(d, 'Đánh số La Mã thường (i, ii, iii, iv...). Đặt ở giữa hoặc bên phải chân trang.')
P(d, 'Phần thân (từ Mở đầu trở đi):', bold=True, indent=False)
P(d, 'Đánh số Ả-rập (1, 2, 3...), đặt ở giữa hoặc bên phải chân trang, cỡ 12 pt.')
P(d, 'Trang bìa và phụ bìa:', bold=True, indent=False)
P(d, 'KHÔNG đánh số.')
d.add_paragraph()

# 6. Tai lieu tham khao
H(d, '6. Tài liệu tham khảo', 2)
add_table(d,
    ['Khối ngành', 'Phong cách trích dẫn phổ biến'],
    [
        ('Khoa học Xã hội — Giáo dục — Tâm lý học', 'APA 7th edition'),
        ('Y Dược', 'Vancouver'),
        ('Khoa học Tự nhiên — Kỹ thuật', 'IEEE hoặc Harvard'),
    ],
    widths=[Cm(8.0), Cm(8.0)])
P(d, 'Quy ước sắp xếp:', bold=True, indent=False)
P(d, '• Tài liệu tiếng Việt trước, tài liệu tiếng Anh/Pháp/khác sau.')
P(d, '• Tiếng Việt: sắp xếp theo abc của TÊN (chữ cuối). Ví dụ: "Công Thị Hằng" sắp theo chữ "H".')
P(d, '• Tiếng Anh: sắp xếp theo abc của HỌ (chữ đầu). Ví dụ: "Anderson, T. L."')
P(d, '• Cỡ chữ: 13 pt, giãn dòng 1,15 (compact). Thụt treo (hanging indent) 1,0 cm để dễ nhìn từng mục.')
d.add_paragraph()

# 7. Do dai
H(d, '7. Độ dài', 2)
add_table(d,
    ['Phần', 'Độ dài'],
    [
        ('Luận án TS (phần thân)', '150–250 trang (không kể TLTK + phụ lục)'),
        ('Tóm tắt luận án (tiếng Việt)', '24 trang'),
        ('Tóm tắt luận án (tiếng Anh)', '24 trang'),
        ('Mở đầu', '5–7 trang'),
        ('Mỗi chương lý luận/thực trạng', '30–60 trang'),
    ],
    widths=[Cm(8.0), Cm(8.0)])
d.add_paragraph()

# 8. Cau truc dien hinh
H(d, '8. Cấu trúc điển hình của luận án TS', 2)
items_structure = [
    'Trang bìa cứng (màu xanh đậm hoặc đỏ đậm tùy trường, chữ vàng/nhũ, gáy in dọc).',
    'Trang phụ bìa (giống bìa cứng nhưng in trên giấy thường).',
    'Lời cam đoan của nghiên cứu sinh (NCS ký).',
    'Lời cảm ơn.',
    'Mục lục.',
    'Danh mục các từ viết tắt, ký hiệu.',
    'Danh mục các bảng (Bảng 1.1, 1.2,…).',
    'Danh mục các hình, biểu đồ, sơ đồ.',
    'Mở đầu (lý do chọn đề tài, mục tiêu, nhiệm vụ, đối tượng, khách thể, phạm vi, giả thuyết, ý nghĩa lý luận và thực tiễn, đóng góp mới, cấu trúc luận án).',
    'Chương 1: Cơ sở lý luận / Tổng quan tài liệu.',
    'Chương 2: Tổ chức và phương pháp nghiên cứu.',
    'Chương 3: Kết quả nghiên cứu thực trạng.',
    'Chương 4 (nếu có): Đề xuất biện pháp / Thực nghiệm.',
    'Kết luận và Kiến nghị.',
    'Danh mục công trình của tác giả đã công bố liên quan đến luận án.',
    'Tài liệu tham khảo.',
    'Phụ lục (bảng hỏi, đề cương phỏng vấn, dữ liệu thô, biên bản…).',
]
for i, it in enumerate(items_structure, 1):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'{i:>2}. {it}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.add_paragraph()

# 9. Khuyen nghi
H(d, '9. Khuyến nghị thực tiễn', 2)
P(d, 'a) Đối chiếu quy định của cơ sở đào tạo:', bold=True, indent=False)
P(d, 'Mỗi trường có quy định riêng kèm template Word — thường được công bố trên website Phòng/Khoa Sau đại học hoặc Viện đào tạo SĐH. Trước khi nộp bản chính thức, em sẽ đối chiếu với "Quy định trình bày luận án tiến sĩ" của trường mình để đảm bảo khớp.')
P(d, 'b) Dùng Styles của Word:', bold=True, indent=False)
P(d, 'Nên dùng Styles (Heading 1, Heading 2, Heading 3…) trong Word để tự sinh mục lục — tránh format thủ công gây sai lệch khi in. Khi cần đổi font/cỡ chữ tiêu đề, chỉ cần đổi Style là toàn bộ chương tự cập nhật.')
P(d, 'c) Kiểm tra trước khi in:', bold=True, indent=False)
P(d, 'Trước khi in bìa cứng, in thử 1 bản trên giấy A4 thường để rà soát: tràn lề, sai chương, hình mờ, bảng vỡ, số trang nhảy cách. In bìa cứng tốn 800.000–1.500.000 đồng/bản nên cần kiểm tra kỹ.')
P(d, 'd) Đính kèm bìa CD:', bold=True, indent=False)
P(d, 'Hầu hết các trường yêu cầu nộp kèm 1 đĩa CD chứa file PDF luận án + tóm tắt + danh mục công trình + slide bảo vệ.')
d.add_paragraph()

# 10. Bang tom tat nhanh
H(d, '10. Bảng tóm tắt nhanh các thông số chính', 2)
add_table(d,
    ['Thông số', 'Giá trị khuyến nghị'],
    [
        ('Phông chữ', 'Times New Roman, 13 pt'),
        ('Giãn dòng', '1,5 lines'),
        ('Lề (T/B/T/P)', '2,5 / 2,5 / 3,5 / 2,0 cm'),
        ('Thụt đầu dòng', '1,25 cm'),
        ('Căn lề văn bản', 'Justify (đều hai bên)'),
        ('Khoảng cách sau đoạn', '6 pt (hoặc 0 pt)'),
        ('Đánh số trang', 'La Mã thường cho phần đầu; Ả-rập từ Mở đầu'),
        ('Số trang trên trang', 'Giữa hoặc phải chân trang, 12 pt'),
        ('Tiêu đề chương', '16–18 pt, IN HOA, ĐẬM, căn giữa'),
        ('Tiêu đề mục 1.1', '14 pt, đậm, trái'),
        ('Tiêu đề mục 1.1.1', '13 pt, đậm, trái'),
        ('Bảng/Hình', 'Bảng X.Y trên / Hình X.Y dưới'),
        ('TLTK style', 'APA 7th (KHXH-GD-TLH)'),
        ('Độ dài thân luận án', '150–250 trang'),
        ('Tóm tắt luận án', '24 trang VN + 24 trang EN'),
    ],
    widths=[Cm(7.0), Cm(9.0)])
d.add_paragraph()

P(d, 'Thầy có muốn em soạn riêng một file Template Word đã định dạng sẵn theo bảng quy chuẩn trên để dùng làm khung cho luận án không ạ? Hoặc nếu thầy có quy định cụ thể của trường mình, em sẽ đối chiếu và điều chỉnh template cho khớp.', italic=True)
d.add_paragraph()

P(d, 'Em xin cảm ơn thầy.', indent=False)
P(d, 'Trân trọng,', indent=False)
P(d, '', indent=False)
P(d, 'NCS Công Thị Hằng', bold=True, indent=False)

# Remove watermark + header/footer
for sec in d.sections:
    for hf in [sec.header, sec.first_page_header, sec.even_page_header,
               sec.footer, sec.first_page_footer, sec.even_page_footer]:
        for elem in list(hf._element.iter()):
            if elem.tag in (qn('w:pict'), qn('w:object')):
                if elem.getparent() is not None:
                    elem.getparent().remove(elem)
        for p in hf.paragraphs:
            for r in p.runs:
                r.text = ''

d.save(OUT)
from docx import Document as D
dd = D(OUT)
w = sum(len(p.text.split()) for p in dd.paragraphs)
print(f"Da luu: {OUT}")
print(f"Paragraphs: {len(dd.paragraphs)}, Tables: {len(dd.tables)}, Words: ~{w}, Size: {os.path.getsize(OUT)//1024}KB")
