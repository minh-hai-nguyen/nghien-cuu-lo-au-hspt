# -*- coding: utf-8 -*-
"""Generate: Phân biệt MA vs NMA (Bayesian) — cho thầy, cùng style với KTC v3."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Giai_thich_MA_vs_NMA_cho_thay.docx')

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def PRun(parts):
    p = doc.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 12))
        if opts.get('bold'): r.bold = True
        if opts.get('italic'): r.italic = True
        if opts.get('red'): r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif opts.get('color'): r.font.color.rgb = opts['color']
    return p

def H1(t): return P(t, bold=True, size=16, color=RGBColor(0x1F, 0x3A, 0x68), align=WD_ALIGN_PARAGRAPH.CENTER)
def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10.5)
    return t

def note_box(text, color_hex='FFF3E0'):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]
    c.text = ''
    r = c.paragraphs[0].add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    set_cell_bg(c, color_hex)


# ============================================================
# TIÊU ĐỀ
# ============================================================
H1('PHÂN BIỆT MA vs NMA (BAYESIAN)')
H1('Meta-Analysis và Network Meta-Analysis')
P('Tài liệu đi kèm với "Giải thích KTC vs CrI" — bổ sung cho thầy khi đọc các báo cáo Bayesian',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=RGBColor(0x66, 0x66, 0x66))
P('Nhóm dự án Lo âu | Tháng 04/2026',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P()

# ============================================================
# TÓM TẮT 30 GIÂY
# ============================================================
note_box(
    'TÓM TẮT TRONG 1 CÂU:\n\n'
    'MA = Meta-Analysis = so sánh 2 can thiệp (thường là can thiệp vs đối chứng). '
    'NMA = Network Meta-Analysis = so sánh NHIỀU can thiệp cùng lúc (≥ 3), dùng cả bằng '
    'chứng TRỰC TIẾP và GIÁN TIẾP, có thể xếp hạng bằng SUCRA. Thêm "Bayesian" vào đầu = '
    'dùng framework Bayesian → báo cáo có CrI (Credible Interval) thay vì CI.',
    color_hex='FFFDE7')
P()

# ============================================================
# PHẦN 1 — TỪ VIẾT TẮT
# ============================================================
H2('1. Giải mã các từ viết tắt')

MakeTable(
    ['Viết tắt', 'Tên đầy đủ (EN)', 'Tên tiếng Việt', 'Nghĩa ngắn gọn'],
    [
        ('MA', 'Meta-Analysis', 'Phân tích tổng hợp / Phân tích gộp',
         'Tổng hợp kết quả từ nhiều NC để ước lượng 1 hiệu ứng chung'),
        ('NMA', 'Network Meta-Analysis', 'Phân tích tổng hợp mạng (lưới)',
         'So sánh nhiều can thiệp cùng lúc qua "mạng" bằng chứng'),
        ('SR', 'Systematic Review', 'Tổng quan hệ thống',
         'Tổng hợp ĐỊNH TÍNH (narrative). MA/NMA là định lượng, thường đi kèm sau SR'),
        ('RCT', 'Randomized Controlled Trial', 'Thử nghiệm ngẫu nhiên có đối chứng',
         'Đơn vị cơ bản của MA/NMA (các RCT được pool lại)'),
        ('Bayesian MA', 'Bayesian Meta-Analysis', 'MA theo framework Bayesian',
         'Output: SMD + 95 % CrI (Credible Interval)'),
        ('Bayesian NMA', 'Bayesian Network Meta-Analysis', 'NMA theo framework Bayesian',
         'Output: SMD mọi cặp + CrI + SUCRA xếp hạng'),
        ('SUCRA', 'Surface Under the Cumulative RAnking curve', 'Diện tích dưới đường xếp hạng tích luỹ',
         'Chỉ xuất hiện trong NMA. Từ 0–100 %, càng cao = can thiệp càng hiệu quả tốt'),
    ])
P()

# ============================================================
# PHẦN 2 — MA TRƯỚC (ĐỂ THẦY HIỂU DẦN)
# ============================================================
H2('2. MA (Meta-Analysis) — "tổng hợp 2 nhóm"')

H3('Câu hỏi MA trả lời:')
P('"Tổng hợp tất cả RCT đã so sánh CAN THIỆP A với ĐỐI CHỨNG, hiệu quả TRUNG BÌNH của A là bao nhiêu?"',
  italic=True, color=RGBColor(0x2E, 0x54, 0x8B))
P()

H3('Ví dụ cụ thể: Zhang 2026 (QT049 trong thư viện)')
P('• Câu hỏi: "CBT học đường có hiệu quả giảm lo âu ở HS nguy cơ thấp không?"', size=12)
P('• 31 RCT được tìm thấy — tất cả đều so CBT vs ĐỐI CHỨNG (waitlist hoặc TAU)', size=12)
P('• Mỗi RCT có 1 SMD (Cohen d hoặc Hedges g)', size=12)
P('• MA pool 31 SMD này thành 1 SMD tổng hợp: –0,19 (95 % CrI: –0,22; –0,17)', size=12)
P('• Kết luận: CBT có hiệu quả nhỏ giảm lo âu — hiệu ứng chung qua tất cả 31 RCT', size=12)
P()

H3('Sơ đồ mạng bằng chứng của MA:')
P('CBT ────[31 RCT, n=19.865]──── Đối chứng', align=WD_ALIGN_PARAGRAPH.CENTER, bold=True, size=13)
P('(Chỉ có 1 cạnh trong mạng — đó là MA đơn giản nhất)',
  italic=True, size=10, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))
P()

note_box(
    'ĐẶC ĐIỂM MA: Đơn giản, mạnh, nhưng CHỈ trả lời được câu hỏi "A có tốt hơn đối chứng không?". '
    'KHÔNG thể trả lời "A vs B tốt hơn?" nếu chưa có RCT trực tiếp so A vs B.',
    color_hex='E3F2FD')
P()

# ============================================================
# PHẦN 3 — NMA VÍ DỤ THỰC TẾ
# ============================================================
H2('3. NMA (Network Meta-Analysis) — "so nhiều can thiệp cùng lúc"')

H3('Câu hỏi NMA trả lời:')
P('"Trong 10 can thiệp khác nhau cho lo âu, can thiệp NÀO tốt nhất? Xếp hạng giúp tôi ưu tiên triển khai."',
  italic=True, color=RGBColor(0x2E, 0x54, 0x8B))
P()

H3('Ví dụ cụ thể: Xian, Zhang & Jiang 2024 (QT039 — Bayesian NMA)')
P('• Câu hỏi: "Trong các can thiệp tâm lý cho rối loạn lo âu xã hội (SAD) ở vị thành niên, liệu pháp nào hiệu quả nhất?"', size=12)
P('• Thu thập 30 RCT với 9 liệu pháp khác nhau (iCBT, gCBT, individual CBT, mindfulness, exposure therapy...)', size=12)
P('• Tổng n = 1.547 VTN SAD', size=12)
P('• NMA sử dụng Bayesian framework → posterior cho mỗi so sánh cặp', size=12)
P('• Output:', bold=True, size=12)
P('  – Effect size cho mỗi cặp (ví dụ iCBT vs gCBT, iCBT vs waitlist...)', size=12)
P('  – SUCRA xếp hạng: iCBT 71,2 % (hạng 1), gCBT 68,4 % (hạng 2)...', size=12)
P('  – Xác suất là "best" cho mỗi can thiệp', size=12)
P()

H3('Sơ đồ mạng bằng chứng của NMA:')
P('Một mạng có nhiều cạnh — mỗi cạnh là 1 cặp can thiệp đã có RCT so trực tiếp:',
  italic=True, size=11)
P()

# ASCII art network
P('                  iCBT', align=WD_ALIGN_PARAGRAPH.CENTER, size=11, bold=True, color=RGBColor(0xC0, 0x00, 0x00))
P('                  ╱  │  ╲', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('                 ╱   │   ╲', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('          gCBT ═══ i-CBT ═══ Waitlist', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('           │  ╲    │    ╱  │', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('           │   ╲   │   ╱   │', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('         Mindfulness   TAU', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('           │            │', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('         Exposure ── Đối chứng', align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
P('(Ví dụ minh hoạ, không phải data thực của Xian 2024)',
  italic=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER, color=RGBColor(0x80, 0x80, 0x80))
P()

P('Mỗi ĐƯỜNG trong mạng là 1 cặp can thiệp có ≥ 1 RCT so trực tiếp.', italic=True)
PRun([
    ('Điểm HAY của NMA: ', {'bold': True}),
    ('dù KHÔNG có RCT trực tiếp iCBT vs Exposure, NMA vẫn ước lượng được hiệu quả thông qua ', {}),
    ('mạng', {'italic': True}),
    (' (iCBT ↔ gCBT ↔ Mindfulness ↔ Exposure). Đây gọi là bằng chứng ', {}),
    ('GIÁN TIẾP (indirect evidence)', {'bold': True}),
    ('.', {}),
])
P()

# ============================================================
# PHẦN 4 — BẢNG SO SÁNH MA vs NMA
# ============================================================
H2('4. Bảng so sánh MA vs NMA')

MakeTable(
    ['Tiêu chí', 'MA (Meta-Analysis)', 'NMA (Network Meta-Analysis)'],
    [
        ('Số can thiệp so sánh', '2 (A vs đối chứng)', '≥ 3 (A, B, C, D... cùng lúc)'),
        ('Kiểu bằng chứng', 'Chỉ TRỰC TIẾP (A vs B)',
         'Cả TRỰC TIẾP (A vs B) VÀ GIÁN TIẾP (A vs C qua B)'),
        ('Output chính', '1 effect size tổng hợp (vd SMD = –0,19)',
         'Ma trận effect size cho mọi cặp + SUCRA xếp hạng'),
        ('Xếp hạng can thiệp?', 'KHÔNG (chỉ có A tốt hơn đối chứng hay không)',
         'CÓ — SUCRA từ 0 đến 100 %'),
        ('Giả định cần thiết', 'Homogeneity (các NC giống nhau)',
         'Homogeneity + Transitivity (bằng chứng gián tiếp đáng tin)'),
        ('Độ phức tạp', 'Dễ — có thể tính bằng Excel',
         'Phức tạp — cần R/Stan/WinBUGS, đặc biệt Bayesian'),
        ('Hay dùng nhất khi', 'Muốn biết: "A có hiệu quả không?"',
         'Muốn biết: "Trong nhóm can thiệp, cái nào TỐT NHẤT?"'),
        ('Bayesian advantage',
         'Có CrI, có posterior distribution',
         'Có CrI + SUCRA tính trực tiếp từ MCMC samples (ưu thế lớn)'),
    ])
P()

# ============================================================
# PHẦN 5 — HIỂU SUCRA
# ============================================================
H2('5. Hiểu SUCRA (chỉ có trong NMA)')

PRun([
    ('SUCRA = ', {'bold': True}),
    ('S', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('urface ', {}),
    ('U', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('nder the ', {}),
    ('C', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('umulative ', {}),
    ('RA', {'bold': True, 'color': RGBColor(0xC0, 0x00, 0x00)}),
    ('nking curve', {}),
])
P()

P('Cách tính đơn giản (không cần công thức):', bold=True)
P('Với 9 can thiệp trong mạng, Bayesian NMA chạy MCMC (Markov Chain Monte Carlo) 10.000 lần. '
  'Mỗi lần, nó tính hạng 1 → 9 cho từng can thiệp. Sau 10.000 lần, cộng dồn xem bao nhiêu % thời '
  'gian một can thiệp được xếp hạng cao. Đó là SUCRA.', size=11)
P()

H3('Bảng diễn giải SUCRA:')
MakeTable(
    ['SUCRA', 'Diễn giải', 'Xếp hạng tương đối'],
    [
        ('90–100 %', 'Gần như chắc chắn là tốt NHẤT trong các can thiệp được so sánh', 'Hạng 1 (ưu tiên cao)'),
        ('70–90 %', 'Rất có khả năng xếp hạng CAO', 'Hạng 1–2'),
        ('50–70 %', 'Trung bình — có thể tốt hoặc vừa phải', 'Hạng giữa'),
        ('30–50 %', 'Có thể xếp hạng TRUNG BÌNH–THẤP', 'Hạng giữa–dưới'),
        ('0–30 %', 'Gần như chắc chắn xếp hạng THẤP', 'Hạng cuối (ít ưu tiên)'),
    ])
P()

H3('Ví dụ áp dụng cho Xian 2024 (QT039):')
P('Theo Xian, Zhang & Jiang 2024 — 9 liệu pháp cho SAD VTN:', italic=True)
MakeTable(
    ['Can thiệp', 'SUCRA', 'Diễn giải'],
    [
        ('iCBT', '71,2 %', 'Xếp hạng 1 — iCBT có 71 % xác suất là can thiệp tốt nhất cho SAD VTN'),
        ('gCBT (group CBT)', '68,4 %', 'Xếp hạng 2 — gần sát iCBT'),
        ('...', '...', '(các can thiệp khác xếp sau)'),
    ])
P()

note_box(
    'HIỂU ĐƠN GIẢN VỀ SUCRA: SUCRA 71 % nghĩa là nếu thầy bắt thăm ngẫu nhiên các kịch bản '
    'xếp hạng (theo posterior distribution), iCBT có 71 % xác suất được xếp ĐỨNG ĐẦU. Đây '
    'là con số dễ truyền thông cho nhà hoạch định chính sách: "iCBT là lựa chọn tốt nhất "'
    'với xác suất 71 %".',
    color_hex='E8F5E9')
P()

# ============================================================
# PHẦN 6 — PHÉP ẨN DỤ
# ============================================================
H2('6. Phép ẩn dụ giúp nhớ')

H3('MA — như "So sánh 2 món ăn"')
P('Thầy muốn biết phở có ngon hơn bún không? → Ăn 31 tô phở + 31 tô bún (mỗi tô ở một quán khác nhau) → tính điểm trung bình.', italic=True)
P('Kết luận: "Phở có ngon hơn bún 0,19 điểm trên thang Cohen" → xong.', italic=True)
P()

H3('NMA — như "Xếp hạng 9 quán phở tại Hà Nội"')
P('Thầy muốn biết trong 9 quán phở, quán nào NGON NHẤT?', italic=True)
P('• Không thể ăn tất cả 36 cặp (C(9,2) = 36) — quá nhiều!', size=11)
P('• Nhưng có 30 người đã ăn rồi viết review (30 "RCT"):', size=11)
P('  – Một số người so quán 1 vs quán 2', size=11)
P('  – Một số so quán 2 vs quán 3', size=11)
P('  – Một số so quán 1 vs quán 5...', size=11)
P('• Tạo ra 1 "mạng" so sánh → NMA dùng toàn bộ mạng này để XẾP HẠNG 9 quán.', size=11)
P('• Kết quả: "Quán 3 có SUCRA 85 % — rất có khả năng là quán ngon nhất."', size=11)
P()

note_box(
    'ĐIỂM KHÁC BIỆT CỐT LÕI: MA = "Phở vs bún có khác không?". '
    'NMA = "Trong 9 quán, quán nào ngon nhất? Xếp hạng ra sao?"',
    color_hex='FFF3E0')
P()

# ============================================================
# PHẦN 7 — BAYESIAN VS FREQUENTIST
# ============================================================
H2('7. Thêm chữ "Bayesian" vào đầu — có gì khác?')

P('Cả MA và NMA đều có 2 phiên bản tuỳ framework thống kê:', italic=True)
P()

MakeTable(
    ['Loại', 'Framework', 'Đặc điểm', 'Output chính'],
    [
        ('MA frequentist', 'Tần suất (cổ điển)',
         'Dễ tính, phần mềm phổ biến (RevMan, Stata meta)',
         'SMD + 95 % CI'),
        ('Bayesian MA', 'Bayesian',
         'Phức tạp hơn, cần R/Stan. Cho posterior distribution',
         'SMD + 95 % CrI'),
        ('NMA frequentist', 'Tần suất',
         'Gói netmeta trong R. Tính SUCRA khó hơn',
         'Ma trận SMD + P-score (thay cho SUCRA)'),
        ('Bayesian NMA', 'Bayesian',
         'CHUẨN VÀNG hiện nay cho NMA. R + Stan + MCMC',
         'Ma trận SMD + 95 % CrI + SUCRA dễ tính'),
    ])
P()

note_box(
    'VÌ SAO BAYESIAN NMA PHỔ BIẾN HƠN? (1) SUCRA dễ tính từ MCMC samples; (2) Xử lý '
    '"incoherence" (mâu thuẫn giữa bằng chứng trực tiếp vs gián tiếp) tốt hơn; (3) Cho '
    'phép tính xác suất "là can thiệp tốt nhất" trực tiếp. Khoảng 70 % NMA hiện đại trong '
    'tâm lý học dùng Bayesian framework.',
    color_hex='E3F2FD')
P()

# ============================================================
# PHẦN 8 — CÁC BÀI MA/NMA TRONG THƯ VIỆN DỰ ÁN
# ============================================================
H2('8. Các bài MA / NMA trong thư viện dự án Lo âu')

MakeTable(
    ['Canonical', 'Tác giả / Năm', 'Loại', 'Nội dung chính'],
    [
        ('QT029', 'Li et al. 2025 (BMC Psychiatry)',
         'Bayesian NMA',
         '30 RCT × 9 liệu pháp điều trị lo âu trẻ em. CBT SUCRA 0,66'),
        ('QT039', 'Xian, Zhang & Jiang 2024 (JAD)',
         'Bayesian NMA',
         '30 RCT × 9 liệu pháp cho SAD VTN. iCBT SUCRA 71,2 % (hạng 1)'),
        ('QT040', 'Walder et al. 2025 (JMIR)',
         'Frequentist MA',
         '21 RCT DMHI cho SAD. Pooled Hedges g = 0,878 (guided + SAD-specific)'),
        ('QT044', 'Cao / Cai et al. 2025 (Frontiers)',
         'Frequentist MA',
         'SR+MA school resilience interventions'),
        ('QT048', 'Chen et al. 2025 (JAD)',
         'Frequentist MA',
         '3-level MA 141 NC về yếu tố nguy cơ/bảo vệ anxiety COVID'),
        ('QT049', 'Zhang, Liang & Kang 2026 (JCP)',
         'Bayesian MA',
         '31 RCT CBT học đường. SMD anxiety –0,19 (95 % CrI: –0,22; –0,17)'),
        ('QT050', 'Qiaochu & Wang 2025 (CPP)',
         'Systematic Review (không meta-analysis)',
         '9 RCT mobile CBT — đếm positive/negative, không pool'),
    ])
P()

# ============================================================
# PHẦN 9 — CÁCH ĐỌC BÁO CÁO
# ============================================================
H2('9. Khi đọc báo cáo — kiểm tra 3 điểm')

P('[ ] 1. ', bold=True, size=13)
PRun([
    ('MA hay NMA? ', {'bold': True, 'size': 13}),
    ('Nếu có "xếp hạng" hoặc "SUCRA" → NMA. Nếu chỉ 1 SMD duy nhất → MA.', {'size': 12}),
])

P('[ ] 2. ', bold=True, size=13)
PRun([
    ('Framework là Bayesian hay frequentist? ', {'bold': True, 'size': 13}),
    ('Nhìn phần Methods. Nếu có "Bayesian", "MCMC", "posterior", "credible interval" → Bayesian.', {'size': 12}),
])

P('[ ] 3. ', bold=True, size=13)
PRun([
    ('Heterogeneity (I², τ²) được báo không? ', {'bold': True, 'size': 13}),
    ('I² < 25 % = thấp, OK pool; 25-75 % = trung bình; > 75 % = cao, pooled có thể misleading.', {'size': 12}),
])
P()

# ============================================================
# PHẦN 10 — TÓM TẮT CUỐI
# ============================================================
H2('10. Tóm tắt 3 điều cần nhớ')

P('1. ', bold=True, size=13)
PRun([
    ('MA = so 2 can thiệp, 1 effect size tổng hợp. ', {'bold': True, 'size': 13}),
    ('Dùng khi muốn biết: "A có tốt hơn đối chứng không?"', {'size': 12}),
])
P()

P('2. ', bold=True, size=13)
PRun([
    ('NMA = so nhiều can thiệp cùng lúc (mạng), có SUCRA xếp hạng. ', {'bold': True, 'size': 13}),
    ('Dùng khi muốn biết: "Trong nhóm can thiệp, cái nào TỐT NHẤT?"', {'size': 12}),
])
P()

P('3. ', bold=True, size=13)
PRun([
    ('Thêm "Bayesian" = dùng framework Bayesian. ', {'bold': True, 'size': 13}),
    ('Output có 95 % CrI thay vì 95 % CI. Với NMA, Bayesian phổ biến hơn vì tính SUCRA dễ.', {'size': 12}),
])
P()

note_box(
    'CUỐI CÙNG: Khi đọc báo cáo có ghi "Bayesian NMA" → thầy hiểu ngay đây là nghiên cứu tổng '
    'hợp MẠNG nhiều can thiệp (không chỉ 2), dùng framework Bayesian, và có xếp hạng SUCRA. '
    'Thuật ngữ nghe to tát nhưng bản chất CHỈ là "so sánh nhiều thứ cùng lúc + có xếp hạng".',
    color_hex='E8F5E9')
P()

# ============================================================
# PHỤ LỤC
# ============================================================
P('─' * 70, color=RGBColor(0x99, 0x99, 0x99))
H3('Phụ lục: Nguồn tham khảo')
P('• Borenstein M, Hedges LV, Higgins JPT, Rothstein HR. "Introduction to Meta-Analysis" (2nd ed., 2021, Wiley) — Chương 1–6 cho MA',
  size=10)
P('• Dias S, Ades AE, Welton NJ, Jansen JP, Sutton AJ. "Network Meta-Analysis for Decision-Making" (2018, Wiley) — Sách chuẩn cho Bayesian NMA',
  size=10)
P('• Rouse B, Chaimani A, Li T. "Network meta-analysis: an introduction for clinicians" (2017, Intern Emerg Med 12:103–111) — bài review dễ đọc',
  size=10)
P('• Salanti G et al. "Graphical methods and numerical summaries for presenting results from multiple-treatment meta-analysis: an overview and tutorial" (2011, J Clin Epidemiol 64:163–71) — giải thích SUCRA',
  size=10)
P('• Cochrane Handbook v6.4 (2023) — chương 11 về NMA',
  size=10)
P()

P('Tài liệu này đi kèm với "Giải thích KTC vs CrI" (v3) — cùng phong cách chi tiết + dễ hiểu '
  'cho người thực hành lâm sàng/giáo viên tâm lý học.',
  italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT):,} bytes')
