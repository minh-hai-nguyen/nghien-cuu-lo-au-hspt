# -*- coding: utf-8 -*-
"""
qa_advanced_v4_5methods.py — 5 phuong phap kiem tra chat luong nang cao
  #5 Temporal Knowledge Base
  #3 Ontology + Type System
  #2 Provenance Chain
  #6 Claim Strength Classifier
  #4 Contradiction Detection

Input: canonical_index.json, kg_v2.graphml, 68 summaries, 61 translations, report v5-4
Output: qa_5methods_report.json + console
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import networkx as nx
from docx import Document
from collections import defaultdict

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
KG_FILE = os.path.join(os.path.dirname(__file__), 'kg_data', 'kg_v2.graphml')
IDX_FILE = os.path.join(ROOT, '02_Papers-goc', 'canonical_index.json')
REPORT = os.path.join(ROOT, '01_Bao-cao',
                      'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-4.docx')
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
BD_DIR = os.path.join(ROOT, '03_Ban-dich')

# ============================================================
# SHARED: Load data
# ============================================================
print('='*70)
print('QA ADVANCED v4 — 5 METHODS')
print('='*70)

print('\nLoading data...')
with open(IDX_FILE, encoding='utf-8') as f:
    canon_idx = json.load(f)
G = nx.read_graphml(KG_FILE)

def read_docx(path):
    try:
        d = Document(path)
        lines = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for r in t.rows:
                lines.append(' | '.join(c.text.strip() for c in r.cells))
        return '\n'.join(lines)
    except:
        return ''

def get_cid(fn):
    m = re.match(r'(VN\d{3}|QT\d{3})', fn)
    return m.group(1) if m else None

summaries, translations = {}, {}
for f in sorted(os.listdir(TT_DIR)):
    if f.endswith('.docx') and not f.startswith('~') and not f.startswith('_'):
        cid = get_cid(f)
        if cid: summaries[cid] = read_docx(os.path.join(TT_DIR, f))
for f in sorted(os.listdir(BD_DIR)):
    if f.endswith('.docx') and not f.startswith('~') and not f.startswith('_'):
        cid = get_cid(f)
        if cid: translations[cid] = read_docx(os.path.join(BD_DIR, f))
for subdir in os.listdir(BD_DIR):
    sp = os.path.join(BD_DIR, subdir)
    if os.path.isdir(sp):
        for f in os.listdir(sp):
            cid = get_cid(f)
            if cid and cid not in translations:
                if f.endswith('.docx'):
                    translations[cid] = read_docx(os.path.join(sp, f))
                elif f.endswith('.md'):
                    try:
                        with open(os.path.join(sp, f), encoding='utf-8') as fh:
                            translations[cid] = fh.read()
                    except: pass

report_doc = Document(REPORT)
report_paras = [(i, p.text) for i, p in enumerate(report_doc.paragraphs) if p.text.strip()]

print(f'  Papers: {len(canon_idx)}, KG: {G.number_of_nodes()} nodes')
print(f'  Summaries: {len(summaries)}, Translations: {len(translations)}')
print(f'  Report paragraphs: {len(report_paras)}')

all_issues = []

# ============================================================
# METHOD #5: TEMPORAL KNOWLEDGE BASE
# ============================================================
print('\n' + '='*70)
print('METHOD #5: TEMPORAL KNOWLEDGE BASE')
print('='*70)

# Infer data_year and publication_year from summaries
temporal_db = {}
for cid, txt in summaries.items():
    pub_year = None
    data_year = None
    # Find publication year from descriptor
    desc = canon_idx.get(cid, {}).get('descriptor', '')
    ym = re.findall(r'20[12]\d', desc)
    if ym: pub_year = int(ym[-1])
    # Find data collection years in text
    # Patterns: "2019-2023", "tháng 3-6/2022", "T9-12/2020", "2020-2021"
    data_matches = re.findall(r'(?:tháng|T)\s*\d+[-–/]\d+/?(20[12]\d)', txt)
    data_matches += re.findall(r'(20[12]\d)\s*[-–]\s*(20[12]\d)', txt)
    data_matches_flat = []
    for dm in data_matches:
        if isinstance(dm, tuple):
            data_matches_flat.extend(dm)
        else:
            data_matches_flat.append(dm)
    if data_matches_flat:
        years = [int(y) for y in data_matches_flat if 2005 <= int(y) <= 2026]
        if years:
            data_year = min(years)

    if not pub_year:
        # Try from text
        ym2 = re.findall(r'(?:xuất bản|đăng|published|pub)\D*(20[12]\d)', txt[:500])
        if ym2: pub_year = int(ym2[0])
        elif ym: pub_year = int(ym[-1])

    temporal_db[cid] = {
        'pub_year': pub_year,
        'data_year': data_year,
        'data_age': (2026 - data_year) if data_year else None,
    }

# Check for outdated facts used in report
temporal_issues = []
for pi, ptxt in report_paras:
    cids_in_para = re.findall(r'\[(VN\d{3}|QT\d{3})\]', ptxt)
    for cid in cids_in_para:
        td = temporal_db.get(cid, {})
        if td.get('data_age') and td['data_age'] > 6:
            temporal_issues.append({
                'method': 'temporal',
                'severity': 'info' if td['data_age'] <= 8 else 'warning',
                'paper': cid,
                'data_year': td['data_year'],
                'pub_year': td['pub_year'],
                'data_age': td['data_age'],
                'para': pi,
                'note': f'Data collected {td["data_age"]} years ago ({td["data_year"]})',
            })

# Check for superseded data
# V-NAMHS 2022 (data 2019-2020) is latest national VN survey
# GBD 2021 supersedes GBD 2019
supersede_map = {
    # If a paper uses GBD 2019, flag if GBD 2021 is available
}

print(f'  Papers with temporal data: {sum(1 for v in temporal_db.values() if v["data_year"])}')
print(f'  Temporal issues in report: {len(temporal_issues)}')
for t in temporal_issues[:5]:
    print(f'    [{t["paper"]}] data {t["data_year"]}, age {t["data_age"]}y — P{t["para"]}')
all_issues.extend(temporal_issues)

# ============================================================
# METHOD #3: ONTOLOGY + TYPE SYSTEM
# ============================================================
print('\n' + '='*70)
print('METHOD #3: ONTOLOGY + TYPE SYSTEM')
print('='*70)

# Define ontology
INSTRUMENTS = {
    'GAD-7': 'Anxiety', 'DASS-21': 'Mixed', 'DASS-Y': 'Mixed',
    'DASS-42': 'Mixed', 'PHQ-9': 'Depression', 'CESD-R': 'Depression',
    'CES-D': 'Depression', 'CESD': 'Depression',
    'MHT': 'Anxiety', 'STAI': 'Anxiety', 'SAS': 'Anxiety',
    'SIAS': 'SAD', 'PSQI': 'Sleep', 'ESSA': 'AcademicStress',
    'SDQ': 'General', 'SCAS': 'Anxiety',
}

DESIGN_HIERARCHY = {
    'meta-analysis': 6, 'NMA': 6, 'umbrella': 7,
    'systematic review': 5, 'SR': 5,
    'RCT': 4, 'cluster RCT': 4, 'cluster controlled': 4,
    'quasi-experimental': 3, 'cohort': 3, 'longitudinal': 3,
    'cross-sectional': 2, 'cắt ngang': 2,
    'case series': 1, 'case report': 1,
}

# Tag each paper with design + instrument + outcome
paper_types = {}
for cid, txt in summaries.items():
    txt_lower = txt.lower()
    # Detect design
    design = 'cross-sectional'  # default
    for d, level in sorted(DESIGN_HIERARCHY.items(), key=lambda x: -x[1]):
        if d.lower() in txt_lower:
            design = d
            break

    # Detect instrument
    instruments = []
    for inst, outcome in INSTRUMENTS.items():
        if inst in txt or inst.lower() in txt_lower:
            instruments.append((inst, outcome))

    # Detect primary outcome
    outcomes = set()
    for _, o in instruments:
        if o != 'Mixed':
            outcomes.add(o)
    if not outcomes:
        if 'lo âu' in txt_lower or 'anxiety' in txt_lower:
            outcomes.add('Anxiety')
        if 'trầm cảm' in txt_lower or 'depression' in txt_lower:
            outcomes.add('Depression')

    paper_types[cid] = {
        'design': design,
        'design_level': DESIGN_HIERARCHY.get(design, 2),
        'instruments': instruments,
        'outcomes': list(outcomes),
    }

# Check report for type mismatches
type_issues = []

# Rule 1: Don't compare prevalence from different instruments
for pi, ptxt in report_paras:
    # Find comparisons like "X% ... cao hơn/thấp hơn ... Y%"
    if any(w in ptxt for w in ['cao hơn', 'thấp hơn', 'so với', 'chênh lệch']):
        cids = re.findall(r'\[(VN\d{3}|QT\d{3})\]', ptxt)
        if len(cids) >= 2:
            instruments_used = set()
            for cid in cids:
                pt = paper_types.get(cid, {})
                for inst, _ in pt.get('instruments', []):
                    instruments_used.add(inst)
            if len(instruments_used) >= 2:
                type_issues.append({
                    'method': 'ontology',
                    'severity': 'warning',
                    'rule': 'instrument_mismatch_comparison',
                    'para': pi,
                    'papers': cids,
                    'instruments': list(instruments_used),
                    'note': f'Comparing papers using different instruments: {instruments_used}',
                    'context': ptxt[:120],
                })

# Rule 2: Cross-sectional claim using causal language
causal_words = ['gây ra', 'gây nên', 'dẫn đến', 'làm tăng', 'làm giảm', 'giảm nguy cơ',
                'tăng nguy cơ', 'hiệu quả', 'can thiệp', 'intervention']
association_ok = ['liên quan', 'hiệp hội', 'association', 'tương quan', 'correlation']
for pi, ptxt in report_paras:
    cids = re.findall(r'\[(VN\d{3}|QT\d{3})\]', ptxt)
    for cid in cids:
        pt = paper_types.get(cid, {})
        if pt.get('design_level', 2) <= 2:  # cross-sectional
            for cw in causal_words:
                if cw in ptxt and not any(aw in ptxt for aw in ['cắt ngang', 'cross-section',
                                          'không thể kết luận', 'caveat', 'lưu ý',
                                          'hiệp hội', 'association', 'ứng viên can thiệp']):
                    type_issues.append({
                        'method': 'ontology',
                        'severity': 'warning',
                        'rule': 'causal_from_crosssectional',
                        'para': pi,
                        'paper': cid,
                        'design': pt.get('design'),
                        'causal_word': cw,
                        'note': f'Causal language "{cw}" used with cross-sectional study [{cid}]',
                        'context': ptxt[:120],
                    })
                    break  # one per paragraph

print(f'  Papers typed: {len(paper_types)}')
print(f'  Ontology issues in report: {len(type_issues)}')
for t in type_issues[:10]:
    print(f'    [{t.get("paper","?")}] {t["rule"]}: {t["note"][:80]}')
all_issues.extend(type_issues)

# ============================================================
# METHOD #2: PROVENANCE CHAIN
# ============================================================
print('\n' + '='*70)
print('METHOD #2: PROVENANCE CHAIN')
print('='*70)

NUM_PAT = re.compile(r'(?:OR|AOR|aOR|β|d|g|SMD|SUCRA)\s*=\s*[-−]?([\d]+[,.][\d]+)')

provenance_issues = []
checked = 0
for pi, ptxt in report_paras:
    cids = re.findall(r'\[(VN\d{3}|QT\d{3})\]', ptxt)
    nums = NUM_PAT.findall(ptxt)
    if not cids or not nums:
        continue
    for cid in cids:
        sum_txt = summaries.get(cid, '')
        trans_txt = translations.get(cid, '')
        for num in nums:
            num_comma = num.replace('.', ',')
            num_dot = num.replace(',', '.')
            checked += 1
            # Chain: report → summary → translation
            in_summary = num in sum_txt or num_comma in sum_txt or num_dot in sum_txt
            in_translation = num in trans_txt or num_comma in trans_txt or num_dot in trans_txt

            if not in_summary and not in_translation:
                # Could be cross-ref (number belongs to different paper cited in same para)
                # Check if number is in ANY paper's source
                found_elsewhere = False
                for other_cid in cids:
                    if other_cid == cid:
                        continue
                    ot = summaries.get(other_cid, '') + translations.get(other_cid, '')
                    if num in ot or num_comma in ot or num_dot in ot:
                        found_elsewhere = True
                        break
                if not found_elsewhere:
                    provenance_issues.append({
                        'method': 'provenance',
                        'severity': 'high',
                        'para': pi,
                        'paper': cid,
                        'value': num,
                        'in_summary': in_summary,
                        'in_translation': in_translation,
                        'note': f'Value {num} cited with [{cid}] not found in any source',
                        'context': ptxt[:120],
                    })
            elif in_summary and not in_translation and trans_txt:
                # In summary but not translation — might be cross-ref in summary
                pass  # Already handled by kg_crosscheck_all_v1.py

print(f'  Claims checked: {checked}')
print(f'  Provenance breaks: {len(provenance_issues)}')
for p in provenance_issues[:10]:
    print(f'    P{p["para"]} [{p["paper"]}] value={p["value"]}: {p["note"][:80]}')
all_issues.extend(provenance_issues)

# ============================================================
# METHOD #6: CLAIM STRENGTH CLASSIFIER
# ============================================================
print('\n' + '='*70)
print('METHOD #6: CLAIM STRENGTH CLASSIFIER')
print('='*70)

CAUSAL_PATTERNS = [
    (r'(?:gây ra|gây nên|causes?|leads? to)', 'causal', 4),
    (r'(?:làm tăng|làm giảm|increases?|decreases?|reduces?)\s+(?:nguy cơ|risk|tỷ lệ)',
     'causal_risk', 4),
    (r'(?:hiệu quả|effective|efficacy)\s+(?:can thiệp|intervention|treatment)', 'efficacy', 4),
    (r'(?:giảm|reduce|decrease)\s+\d+\s*%\s+(?:nguy cơ|risk)', 'risk_reduction', 3),
    (r'(?:yếu tố bảo vệ|protective factor)', 'protective', 3),
    (r'(?:yếu tố nguy cơ|risk factor)', 'risk_factor', 2),
    (r'(?:liên quan|associated|correlation|tương quan)', 'association', 2),
    (r'(?:gợi ý|suggest|may|có thể)', 'suggestive', 1),
]

CAVEAT_PATTERNS = [
    r'cắt ngang', r'cross-section', r'không thể kết luận.*nhân quả',
    r'reverse causation', r'confounding', r'hiệp hội.*quan sát',
    r'lưu ý.*phương pháp', r'thiết kế.*không.*RCT',
    r'ứng viên can thiệp', r'cần.*RCT.*xác nhận',
]

strength_issues = []
for pi, ptxt in report_paras:
    cids = re.findall(r'\[(VN\d{3}|QT\d{3})\]', ptxt)
    if not cids:
        continue

    # Detect claim strength
    claim_level = 1  # default: suggestive
    claim_type = 'suggestive'
    for pat, ctype, level in CAUSAL_PATTERNS:
        if re.search(pat, ptxt, re.IGNORECASE):
            if level > claim_level:
                claim_level = level
                claim_type = ctype
            break

    # Check if caveat present
    has_caveat = any(re.search(cp, ptxt, re.IGNORECASE) for cp in CAVEAT_PATTERNS)

    # Get max design level of cited papers
    max_design = 0
    for cid in cids:
        pt = paper_types.get(cid, {})
        dl = pt.get('design_level', 2)
        if dl > max_design:
            max_design = dl

    # Flag if claim strength exceeds evidence
    if claim_level >= 3 and max_design <= 2 and not has_caveat:
        strength_issues.append({
            'method': 'claim_strength',
            'severity': 'warning',
            'para': pi,
            'papers': cids,
            'claim_type': claim_type,
            'claim_level': claim_level,
            'max_design_level': max_design,
            'has_caveat': has_caveat,
            'note': f'Claim strength "{claim_type}" (level {claim_level}) exceeds evidence '
                    f'(design level {max_design}) without caveat',
            'context': ptxt[:150],
        })

print(f'  Strength issues: {len(strength_issues)}')
for s in strength_issues[:10]:
    print(f'    P{s["para"]} {s["claim_type"]} (L{s["claim_level"]}) vs design L{s["max_design_level"]}')
    print(f'      {s["context"][:100]}...')
all_issues.extend(strength_issues)

# ============================================================
# METHOD #4: CONTRADICTION DETECTION
# ============================================================
print('\n' + '='*70)
print('METHOD #4: CONTRADICTION DETECTION')
print('='*70)

# Extract (variable, direction, outcome) tuples per paper
DIRECTION_PATTERNS = [
    # (pattern, variable_group, direction)
    (r'(?:nữ|female).*(?:cao hơn|higher|tăng|increase)', 'gender_female', 'risk'),
    (r'(?:nam|male).*(?:cao hơn|higher|tăng|increase).*(?:lo âu|anxiety)', 'gender_male', 'risk'),
    (r'(?:nam|male).*(?:OR|aOR)\s*=\s*0[,.]', 'gender_male', 'protective'),
    (r'(?:áp lực|academic|học tập).*(?:OR|tăng|risk)', 'academic_pressure', 'risk'),
    (r'(?:giấc ngủ|sleep|ngủ).*(?:OR|AOR|tăng|risk|nguy cơ)', 'sleep_deficit', 'risk'),
    (r'(?:gia đình|family|cha mẹ|parent).*(?:bảo vệ|protective|OR\s*=\s*0)', 'family', 'protective'),
    (r'(?:gia đình|family).*(?:nguy cơ|risk|tăng|OR\s*>\s*1)', 'family_conflict', 'risk'),
    (r'(?:vận động|exercise|physical|thể chất).*(?:bảo vệ|protective|giảm)', 'exercise', 'protective'),
    (r'(?:MXH|social media|screen|màn hình).*(?:tăng|risk|nguy cơ|harmful)', 'social_media', 'risk'),
    (r'(?:MXH|social media|internet).*(?:bảo vệ|protective|giảm)', 'social_media', 'protective'),
    (r'(?:cô đơn|lonely|loneliness).*(?:tăng|risk|nguy cơ)', 'loneliness', 'risk'),
    (r'(?:mindfulness).*(?:thất bại|fail|không hiệu quả)', 'mindfulness', 'ineffective'),
    (r'(?:mindfulness).*(?:hiệu quả|effective|giảm)', 'mindfulness', 'effective'),
    (r'(?:CBT).*(?:hiệu quả|effective)', 'CBT', 'effective'),
]

paper_claims = {}
for cid, txt in summaries.items():
    claims = []
    for pat, var, direction in DIRECTION_PATTERNS:
        if re.search(pat, txt, re.IGNORECASE):
            claims.append((var, direction))
    paper_claims[cid] = list(set(claims))

# Detect contradictions
contradictions = []
papers_list = sorted(paper_claims.keys())
for i, p1 in enumerate(papers_list):
    for p2 in papers_list[i+1:]:
        for (v1, d1) in paper_claims[p1]:
            for (v2, d2) in paper_claims[p2]:
                if v1 == v2 and d1 != d2:
                    contradictions.append({
                        'method': 'contradiction',
                        'severity': 'info',
                        'paper1': p1,
                        'paper2': p2,
                        'variable': v1,
                        'direction1': d1,
                        'direction2': d2,
                        'note': f'{p1} says {v1}={d1}, {p2} says {v1}={d2}',
                    })

# Deduplicate
seen = set()
unique_contradictions = []
for c in contradictions:
    key = (c['variable'], frozenset([c['paper1'], c['paper2']]))
    if key not in seen:
        seen.add(key)
        unique_contradictions.append(c)

print(f'  Papers with extracted claims: {sum(1 for v in paper_claims.values() if v)}')
print(f'  Total claim tuples: {sum(len(v) for v in paper_claims.values())}')
print(f'  Contradictions found: {len(unique_contradictions)}')
for c in unique_contradictions:
    print(f'    {c["paper1"]} vs {c["paper2"]}: {c["variable"]} ({c["direction1"]} vs {c["direction2"]})')
all_issues.extend(unique_contradictions)

# ============================================================
# FINAL REPORT
# ============================================================
print('\n' + '='*70)
print('FINAL SUMMARY — 5 METHODS')
print('='*70)

by_method = defaultdict(list)
for iss in all_issues:
    by_method[iss['method']].append(iss)

for method_name, method_label in [
    ('temporal', '#5 Temporal KB'),
    ('ontology', '#3 Ontology + Type System'),
    ('provenance', '#2 Provenance Chain'),
    ('claim_strength', '#6 Claim Strength Classifier'),
    ('contradiction', '#4 Contradiction Detection'),
]:
    issues = by_method.get(method_name, [])
    sevs = defaultdict(int)
    for i in issues:
        sevs[i.get('severity', '?')] += 1
    sev_str = ', '.join(f'{k}:{v}' for k, v in sorted(sevs.items()))
    print(f'  {method_label}: {len(issues)} issues ({sev_str})')

print(f'\n  TOTAL: {len(all_issues)} issues')

# Save JSON
out_path = os.path.join(os.path.dirname(__file__), 'qa_5methods_report.json')
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump({
        'report_file': os.path.basename(REPORT),
        'total_issues': len(all_issues),
        'by_method': {k: len(v) for k, v in by_method.items()},
        'temporal_db_sample': {k: v for k, v in list(temporal_db.items())[:5]},
        'paper_types_sample': {k: v for k, v in list(paper_types.items())[:5]},
        'contradictions': unique_contradictions,
        'issues': all_issues,
    }, f, ensure_ascii=False, indent=2)

print(f'\nFull report: {out_path}')
print('Done.')
