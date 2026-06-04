"""DOC: Tong hop tai lieu VN ve Ap luc hoc tap (ALHT) lan lo au o HS THPT.

ALL FACTS VERIFIED tu Tom-tat-tung-bai (cac bai VN trong canonical):
- VN004 Nguyen Thi Van (2020) STAI TPHCM
- VN006 Tran Thi My Luong (2020) DASS-42 THPT chuyen
- VN020 Tran Ho Vinh Loc (2024) DASS-Y + ESSA TPHCM
- VN025 Pham Thi Ngoc (2024) DASS-21 Hai Phong
- VN029 Truc Thanh Thai (2025) DASS-21 TPHCM Q1 Soc Psychiatry
- VN021 Tran Thao Vi (2024) ESSA Hue (THCS, lam tham chieu)
- VN027 Dinh (2021) yeu to truong hoc THCS
- VN015 Ngo Anh Vinh (2024) DTTS Lang Son (lam tham chieu)
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Tong_hop_tai_lieu_VN_ALHT_lo_au_HS_THPT.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('TỔNG HỢP TÀI LIỆU VIỆT NAM\nVỀ ÁP LỰC HỌC TẬP VỚI LO ÂU\nỞ HỌC SINH TRUNG HỌC PHỔ THÔNG\n— 8 nghiên cứu chính (2020–2025) —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Bạn tìm tài liệu về áp lực học tập lên lo âu ở học sinh THPT '
    'ở VN nhé.', italic=True
)

# Bảng tổng quan
H('Bảng 1. Tổng quan 8 nghiên cứu Việt Nam (5 trực tiếp HS THPT + 3 tham chiếu)', level=2, color=NAVY)
add_table(
    ['ID', 'Tác giả (năm)', 'Mẫu', 'Cấp học', 'Thang đo', 'Tỷ lệ lo âu', 'Đo ALHT?'],
    [
        ['VN004', 'Nguyễn Thị Vân (2020)',
         'n=558→90 TPHCM', 'THPT', 'STAI + tự thiết kế',
         '15–18,5%', 'Bảng hỏi tự thiết kế (r=0,37)'],
        ['VN006', 'Trần Thị Mỵ Lương (2020)',
         'n=540 THPT chuyên', 'THPT chuyên', 'DASS-42',
         '14,2%', 'KHÔNG đo trực tiếp'],
        ['VN020', 'Trần Hồ Vĩnh Lộc và cộng sự (2024)',
         'n=976 TPHCM', 'THPT', 'DASS-Y + ESSA',
         '25,1%', '✓ ESSA chuẩn — yếu tố MẠNH NHẤT'],
        ['VN025', 'Phạm Thị Ngọc và cộng sự (2024)',
         'n=420 Hải Phòng', 'THPT + GDNN-GDTX', 'DASS-21',
         '53,81% / 43,33%', 'Bảng hỏi áp lực (yếu tố ý nghĩa)'],
        ['VN029', 'Trúc Thanh Thái và cộng sự (2025)',
         'n=2.631 TPHCM', 'THPT (chính quy + GDTX)', 'DASS-21',
         '50,3%', 'KHÔNG đo trực tiếp ALHT'],
        ['VN021', 'Trần Thảo Vi và cộng sự (2024)',
         'n=341 Huế', 'THCS (tham chiếu dọc)', 'ESSA chuẩn',
         '—', '✓ ESSA full 16 mục dọc 3 năm'],
        ['VN027', 'Đinh và cộng sự (2021)',
         'HS THCS VN', 'THCS (tham chiếu)', 'Tự thiết kế',
         '—', 'ALHT có ý nghĩa'],
        ['VN015', 'Ngô Anh Vinh và cộng sự (2024)',
         'n=845 Lạng Sơn DTTS', 'DTTS nội trú', 'DASS-21',
         '54,4%', 'KHÔNG đo trực tiếp'],
    ]
)
para('')

# I. Năm bài chính HS THPT
H('I. NĂM NGHIÊN CỨU CHÍNH — HS THPT trực tiếp', level=2, color=NAVY)

H('I.1. VN020 Trần Hồ Vĩnh Lộc và cộng sự (2024) — TỐT NHẤT về phương pháp', level=3, color=NAVY)
para(
    'Tác giả: Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh, Tô Gia Kiên (Khoa Y tế Công '
    'cộng, Đại học Y Dược TPHCM, 2024). Tạp chí Y học TPHCM.'
)
para(
    'Mẫu: n = 976 HS THPT tại 2 trường TPHCM. Lấy mẫu ngẫu nhiên nhiều bậc.'
)
para(
    'Công cụ: DASS-Y (Depression Anxiety Stress Scales for Youth, Szabó 2010) '
    '— phiên bản DÀNH RIÊNG cho thanh thiếu niên với ngưỡng cắt riêng (lo âu '
    '≥6, trầm cảm ≥7, stress ≥12 — thấp hơn DASS-21 người lớn). KẾT HỢP với '
    'ESSA (Educational Stress Scale for Adolescents, Sun và cộng sự 2011) gồm '
    '16 mục, 5 tiểu thang, α = 0,81 — đo CHUẨN ÁP LỰC HỌC TẬP.', bold=True
)
para('Phát hiện chính:', bold=True)
para('• Tỷ lệ lo âu (DASS-Y): 25,1% — THẤP HƠN đáng kể so với các NC dùng DASS-21 (40–86%).')
para('• ÁP LỰC HỌC TẬP (ESSA ≥ 59) là YẾU TỐ MẠNH NHẤT — tăng nguy cơ cả 3 chỉ số DASS-Y.')
para('• Cha mẹ ly hôn/ly thân tăng nguy cơ trầm cảm; kinh tế khó khăn tăng trầm cảm + căng thẳng.')
para('• Nữ > nam ở cả 3 chỉ số (p < 0,05) — phù hợp xu hướng toàn cầu.')
para(
    'Đánh giá: ⭐⭐⭐ Đại học Y Dược TPHCM uy tín; DASS-Y phiên bản VTN '
    'riêng (đóng góp phương pháp luận); ESSA chuẩn hóa quốc tế. Hạn chế: '
    'chỉ 2 trường, chọn lớp thuận tiện, tạp chí trong nước không PubMed.'
)
para(
    'GIÁ TRỊ ĐẶC BIỆT: VN020 là MỘT TRONG SỐ ÍT bài VN dùng đồng thời DASS-Y '
    '(thang chuyên VTN) + ESSA (thang chuẩn quốc tế cho ALHT) — phù hợp '
    'chuẩn quốc tế Pascoe (2020) và Sun (2011).', bold=True
)

H('I.2. VN025 Phạm Thị Ngọc và cộng sự (2024) — Đối chiếu THPT chính quy vs GDNN-GDTX', level=3, color=NAVY)
para(
    'Tác giả: Phạm Thị Ngọc, Hoàng Thị Giang, Phạm Khánh Linh, Vũ Thị Châu '
    '(2024). Tạp chí Y học Dự phòng Việt Nam, tập 34, số 1 Phụ bản.'
)
para(
    'Mẫu: n = 420 HS từ 2 cơ sở giáo dục huyện Vĩnh Bảo, Hải Phòng — bao gồm '
    'cả THPT chính quy (Cộng Hiền) + GDNN-GDTX (đối tượng ít được nghiên cứu).'
)
para(
    'Công cụ: DASS-21 phiên bản tiếng Việt + bảng hỏi nhân khẩu học và yếu '
    'tố nguy cơ (giới tính, quan hệ bạn bè, ÁP LỰC HỌC TẬP VÀ THI CỬ, kỳ '
    'vọng gia đình, mâu thuẫn gia đình, bạo lực học đường).'
)
para('Phát hiện chính:', bold=True)
para('• Tỷ lệ lo âu THPT Cộng Hiền: 53,81%; GDNN-GDTX: 43,33% — chênh ~10 điểm %.')
para('• Áp lực học tập VÀ THI CỬ là YẾU TỐ NGUY CƠ ý nghĩa cho LO ÂU.')
para('• Bạo lực học đường: yếu tố nguy cơ độc lập cho lo âu.')
para('• Quan hệ không hòa đồng với bạn bè: yếu tố nguy cơ cho cả trầm cảm và lo âu.')
para('• Cộng Hiền (chính quy) có tỷ lệ CAO HƠN GDNN-GDTX — gợi ý áp lực học thuật đóng vai trò quan trọng.')
para(
    'Đánh giá: ⭐⭐⭐ Tạp chí Y học Dự phòng VN; mẫu vừa; ưu điểm so sánh '
    '2 cơ sở (THPT vs GDNN-GDTX). Hạn chế: chỉ 1 huyện, mẫu thuận tiện, '
    'cắt ngang.'
)

H('I.3. VN029 Trúc Thanh Thái và cộng sự (2025) — Mẫu LỚN NHẤT, Q1 quốc tế', level=3, color=NAVY)
para(
    'Tác giả: Trúc Thanh Thái, Minh Cường Dương và cộng sự (2025). Tạp chí '
    'Social Psychiatry and Psychiatric Epidemiology (Q1 quốc tế, peer-reviewed).'
)
para(
    'Mẫu: n = 2.631 HS THPT TPHCM tại 8 cơ sở giáo dục (4 trường THPT chính '
    'quy + 4 trung tâm GDTX) — MẪU LỚN NHẤT trong các NC SKTT HS THPT VN '
    'có chỉ mục quốc tế.'
)
para(
    'Công cụ: DASS-21 sàng lọc trầm cảm/lo âu/căng thẳng + YBRS đo 8 nhóm '
    'hành vi nguy cơ (sử dụng chất, bạo lực, hành vi nguy cơ tình dục, ...). '
    'Khung Jessor về liên kết hành vi vấn đề.'
)
para('Phát hiện chính:', bold=True)
para('• Tỷ lệ lo âu (DASS-21): 50,3% — phù hợp xu hướng VN dùng DASS-21 (40–86%).')
para('• HS có triệu chứng SOMD (stress/anxiety/mood disorder) có nguy cơ HRB cao hơn 1,24–4,64 lần — bằng chứng MẠNH cho cơ chế "tự dùng".')
para('• KHÔNG đo trực tiếp ALHT — đây là HẠN CHẾ.')
para(
    'Đánh giá: ⭐⭐⭐⭐⭐ Q1 Social Psychiatry; n = 2.631; đa trung tâm bao '
    'phủ cả THPT chính quy + GDTX; đo đồng thời SOMD và HRB. Hạn chế: '
    'cắt ngang, chỉ TPHCM, KHÔNG đo ALHT chuyên biệt.'
)
para(
    'GIÁ TRỊ: VN029 là BÀI VN TỐT NHẤT về dữ liệu VTN TPHCM hiện có; có thể '
    'trích cho TỶ LỆ + giới tính + GDTX vs chính quy. KHÔNG nên trích cho '
    'ALHT vì không đo trực tiếp.', bold=True
)

H('I.4. VN006 Trần Thị Mỵ Lương (2020) — HS THPT chuyên, phát hiện đặc biệt', level=3, color=NAVY)
para(
    'Mẫu: n = 540 HS THPT chuyên — 77 HS có RLLA (tỷ lệ 14,2%). Sử dụng '
    'DASS-42 (40 mục về lo âu).'
)
para('Phát hiện chính:', bold=True)
para('• Tỷ lệ lo âu tổng: 14,2% (nhẹ 3,5% / vừa 7,2% / nặng 2,4% / rất nặng 1,1%).')
para('• PHÁT HIỆN ĐẶC BIỆT: Khối 11 chiếm 48,1% (CAO NHẤT); khối 10 31,1%; khối 12 chỉ 20,8%.')
para(
    '• TRÁI giả thuyết "thi ĐH gây lo âu nhất" — khối 12 thấp nhất. Lý '
    'giải: khối 11 áp lực đa chiều chọn ngành/khối; khối 12 đã ổn định '
    'định hướng. → "Perfectionism trap" trường chuyên.'
)
para('• Lệch giới: nữ 66,5%.')
para(
    'Đánh giá: ⭐⭐⭐⭐ Phát hiện đặc sắc cho HS chuyên; DASS-42 full version. '
    'Hạn chế: chỉ 1 trường chuyên đặc thù; KHÔNG đo trực tiếp ALHT (chỉ '
    'phân tầng theo khối).'
)
para(
    'GIÁ TRỊ: VN006 cho thấy ALHT ở HS THPT chuyên VN có pattern KHÁC HS '
    'thường — khối 11 áp lực cao hơn khối 12. Phù hợp khi thầy bàn về '
    'HS chuyên.', bold=True
)

H('I.5. VN004 Nguyễn Thị Vân (2020) — đã có doc phản biện riêng', level=3, color=NAVY)
para(
    'Tóm tắt nhanh: n = 558 sàng lọc → 90 phỏng vấn sâu HS THPT TPHCM. STAI + '
    'bảng hỏi tự thiết kế. Tỷ lệ RLLA 15–18,5%. Nhóm "học tập" có r = 0,37 '
    'với lo âu (mạnh thứ 2 sau "bản thân HS"). Top 3 biểu hiện: áp lực thi '
    'ĐH 56,7%; định hướng nghề 51,5%; kỳ vọng cha mẹ 48,9%.'
)
para(
    'HẠN CHẾ NGHIÊM TRỌNG: dùng bảng hỏi tự thiết kế chỉ MỘT biến tổng — '
    'không phân biệt 6 chiều con của áp lực học tập (đã phản biện chi tiết '
    'trong doc Phan_bien_VN004_yeu_to_hoc_tap.docx).'
)

# II. Bài tham chiếu
H('II. BA NGHIÊN CỨU THAM CHIẾU (HS THCS / DTTS)', level=2, color=NAVY)

H('II.1. VN021 Trần Thảo Vi và cộng sự (2024) — DỌC 3 NĂM HS THCS Huế', level=3, color=NAVY)
para(
    'Tác giả: Thao Vi Tran, Hoang Thuy Linh Nguyen, Xuan Minh Tri Tran, Yuri '
    'Tashiro, Kaoruko Seino, Thang Van Vo, Keiko Nakamura (Tokyo Medical and '
    'Dental University + ĐH Y Dược Huế). Journal of Rural Medicine, 2024 Oct; '
    '19(4):279–290. DOI: 10.2185/jrm.2024-012.'
)
para(
    'Mẫu: n = 341 HS THCS Huế (thuần tập Hue Healthy Adolescent Cohort '
    '2018–2021). Đây là NGHIÊN CỨU DỌC ĐẦU TIÊN tại VN cho HS THCS về căng '
    'thẳng học tập.'
)
para(
    'Công cụ: ESSA full 16 mục (Sun và cộng sự 2011), 5 lĩnh vực, α = 0,88. '
    'Hồi quy tuyến tính đa biến.'
)
para('Phát hiện chính:', bold=True)
para('• ESSA tăng đáng kể từ M = 46,4 (lớp 6) → 53,5 (lớp 9) — TĂNG 15,3% trong 3 năm.')
para('• 4/5 lĩnh vực ESSA TĂNG; chỉ "lo lắng về điểm số" GIẢM (duy nhất).')
para('• Học thêm là yếu tố MẠNH NHẤT (β = 4,73). 92,1% HS học thêm — TỶ LỆ CỰC CAO.')
para('• Cha học vấn cao → con stress cao (β = 3,20). Mẹ KHÔNG ảnh hưởng.')
para('• Nhiều anh chị em (β = 2,24).')
para(
    'GIÁ TRỊ: VN021 cung cấp BẰNG CHỨNG DỌC duy nhất tại VN về ALHT — phù hợp '
    'tham chiếu cho luận án HS THCS. Tuy nhiên là HS THCS (không phải THPT) — '
    'cần lưu ý khi đối chiếu với HS THPT.', bold=True
)

H('II.2. VN027 Đinh và cộng sự (2021) — Yếu tố trường học', level=3, color=NAVY)
para(
    'Khảo sát yếu tố trường học liên quan lo âu ở HS THCS Việt Nam. Phát hiện '
    'các yếu tố có ý nghĩa: ÁP LỰC HỌC TẬP, kỳ vọng từ giáo viên, mối quan '
    'hệ tiêu cực với bạn bè (bắt nạt), khối lượng bài tập, ít thời gian nghỉ.'
)
para(
    'GIÁ TRỊ: bài VN cụ thể về YẾU TỐ TRƯỜNG HỌC — bổ sung cho UNICEF VN022 '
    '(School Factors). Hạn chế: ResearchGate (preprint), không phải tạp chí '
    'Q1, không có thông tin số liệu chi tiết β/r.'
)

H('II.3. VN015 Ngô Anh Vinh và cộng sự (2024) — DTTS nội trú Lạng Sơn', level=3, color=NAVY)
para(
    'Mẫu: n = 845 HS dân tộc thiểu số nội trú Lạng Sơn. DASS-21 + ACEs.'
)
para('• Tỷ lệ lo âu: 54,4% — RẤT CAO.')
para('• ACEs (Adverse Childhood Experiences) liên quan dương tính với cả 3 chỉ số SKTT (Coef. = 0,16–0,28).')
para('• Bạn bè kém tăng nguy cơ trầm cảm cực mạnh (OR = 6,84).')
para('• KHÔNG đo trực tiếp ALHT — chỉ đo ACEs + tình bạn + internet.')
para(
    'GIÁ TRỊ: VN015 cho thấy DTTS nội trú có tỷ lệ lo âu RẤT CAO. KHÔNG dùng '
    'cho phân tích ALHT.'
)

# III. Tổng kết
H('III. Tổng kết — Bảng so sánh chi tiết', level=2, color=NAVY)
caption('Bảng 2. So sánh phương pháp đo ALHT trong 5 bài chính HS THPT')
add_table(
    ['Bài', 'Đo ALHT?', 'Thang đo cụ thể', 'Số chiều', 'Hệ số đo lường'],
    [
        ['VN004 (2020)', 'Bảng tự thiết kế', 'Không tên thang chuẩn', '1 biến tổng',
         'r = 0,37 với lo âu'],
        ['VN006 (2020)', 'KHÔNG (gián tiếp qua khối lớp)',
         '—', '—', '—'],
        ['VN020 (2024)', '✓ ESSA chuẩn quốc tế (Sun 2011)',
         '16 mục, α=0,81', '5 chiều',
         'ESSA ≥ 59 yếu tố mạnh nhất'],
        ['VN025 (2024)', 'Bảng tự thiết kế (áp lực HT + thi cử)',
         'Không tên thang chuẩn', 'Không công bố',
         'Yếu tố ý nghĩa thống kê'],
        ['VN029 (2025)', 'KHÔNG đo trực tiếp',
         '—', '—', '—'],
    ]
)
para('')
para(
    'Kết luận: Trong 5 bài VN trực tiếp HS THPT, CHỈ VN020 dùng thang ESSA '
    'chuẩn quốc tế. VN004 và VN025 dùng bảng tự thiết kế (hạn chế phương '
    'pháp luận). VN006 và VN029 KHÔNG đo trực tiếp ALHT. Phù hợp với '
    'KHOẢNG TRỐNG mà đề tài của thầy có thể lấp.', bold=True
)

# IV. Khuyến nghị
H('IV. KHUYẾN NGHỊ cho thầy', level=2, color=NAVY)
blue_run('Bốn khuyến nghị khi trích các bài này vào báo cáo CTH:', bold=True)
blue_run(
    '(1) ƯU TIÊN trích VN020 (Trần Hồ Vĩnh Lộc 2024) làm bằng chứng VN '
    'CHẤT LƯỢNG TỐT nhất về ALHT ở HS THPT — vì dùng thang ESSA chuẩn '
    'quốc tế. Trích nguyên: "ESSA ≥ 59 là yếu tố mạnh nhất tăng nguy cơ '
    'lo âu cả 3 chỉ số DASS-Y" trên 976 HS THPT TPHCM.'
)
blue_run(
    '(2) Bổ sung VN025 (Phạm Thị Ngọc 2024) cho phân biệt THPT chính '
    'quy vs GDNN-GDTX — phát hiện chính quy (53,81%) > GDTX (43,33%) gợi '
    'ý áp lực học thuật chính quy mạnh hơn.'
)
blue_run(
    '(3) Trích VN029 (Trúc Thanh Thái 2025 Q1) cho TỶ LỆ lo âu HS THPT '
    'TPHCM (50,3%, n=2.631) — bài VN có chỉ mục quốc tế tốt nhất hiện '
    'có. KHÔNG trích cho ALHT vì không đo.'
)
blue_run(
    '(4) Trích VN021 (Trần Thảo Vi 2024 J Rural Medicine) làm THAM CHIẾU '
    'DỌC duy nhất tại VN — ESSA tăng 15,3% qua 3 năm THCS Huế. Lưu ý '
    'là HS THCS, không phải THPT — chỉ dùng làm bằng chứng "ALHT tăng '
    'theo khối lớp".'
)

blue_run('Năm khoảng trống mà đề tài CTH có thể lấp:', bold=True)
blue_run(
    '(1) Hầu hết NC VN dùng bảng tự thiết kế thay vì ESSA chuẩn — '
    'CTH dùng ESSA rút gọn 4 mục (đã tốt hơn).'
)
blue_run(
    '(2) Hầu hết NC chỉ một thời điểm — CTH có thể bổ sung NC dọc '
    '(theo VN021).'
)
blue_run(
    '(3) Không bài VN nào dùng SEM tích hợp YTNC + YTBV — chương 3 '
    'luận án CTH là ĐẦU TIÊN làm điều này.'
)
blue_run(
    '(4) Không bài VN nào phân tích β chuẩn hóa cho từng dạng RLLA '
    '(RLLATQ/RLLAXH/RLLACL) — CTH có lợi thế.'
)
blue_run(
    '(5) Không bài VN nào tính Cohen d cho hiệu ứng giới — CTH có '
    'thể bổ sung (xem doc Kiem_tra_chenh_lech_gioi đã làm).'
)

# V. Phụ lục TLTK
H('V. Phụ lục — Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Đinh, V. T., và cộng sự. (2021). School factors causing Vietnamese adolescents anxiety. ResearchGate. [VN027 trong cơ sở dữ liệu dự án.]',
    'Hoàng, T. H., & Nguyễn, T. D. (2025). Mức độ căng thẳng, lo âu và trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 tại Việt Nam. [VN014 trong cơ sở dữ liệu — KHÔNG đo ALHT.]',
    'Ngô, A. V., và cộng sự. (2024). Mental Health among Ethnic Minority Adolescents in Vietnam and Correlated Factors. [VN015 trong cơ sở dữ liệu dự án.]',
    'Nguyễn, T. V. (2020). Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh. [VN004 trong cơ sở dữ liệu dự án.]',
    'Phạm, T. N., Hoàng, T. G., Phạm, K. L., & Vũ, T. C. (2024). Thực trạng sức khoẻ tâm thần và một số yếu tố liên quan của học sinh tại 2 trường trung học phổ thông huyện Vĩnh Bảo, Hải Phòng năm 2023. Tạp chí Y học Dự phòng Việt Nam, 34(1 Phụ bản). [VN025.]',
    'Sun, J., Dunne, M. P., Hou, X. Y., & Xu, A. Q. (2011). Educational Stress Scale for Adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546.',
    'Trần, H. V. L., Huỳnh, N. V. A., & Tô, G. K. (2024). Trầm cảm, lo âu, căng thẳng và các yếu tố liên quan ở học sinh THPT tại Thành phố Hồ Chí Minh. Tạp chí Y học TPHCM. [VN020.]',
    'Trần, T. M. L. (2020). Thực trạng rối loạn lo âu ở học sinh trung học phổ thông chuyên. [VN006 trong cơ sở dữ liệu dự án.]',
    'Trúc, T. T., Dương, M. C., và cộng sự. (2025). Unmasking the burden of mental health symptoms and risk behaviors in Vietnamese adolescents: evidence from a multicenter cross-sectional study involving 2,631 high school students. Social Psychiatry and Psychiatric Epidemiology. [VN029.]',
    'Tran, T. V., Nguyen, H. T. L., Tran, X. M. T., Tashiro, Y., Seino, K., Vo, T. V., & Nakamura, K. (2024). Academic stress among students in Vietnam: A three-year longitudinal study. Journal of Rural Medicine, 19(4), 279–290. https://doi.org/10.2185/jrm.2024-012 [VN021.]',
    'Szabó, M. (2010). The short version of the Depression Anxiety Stress Scales (DASS-21): Factor structure in a young adolescent sample. Journal of Adolescence.',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
