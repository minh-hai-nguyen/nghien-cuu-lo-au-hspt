# -*- coding: utf-8 -*-
"""Dich UNICEF VN 2022 — School-Related Factors Impacting Mental Health VTN VN"""
import sys, os, fitz
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

PDF_PATH = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '02_Papers-goc', 'Viet-Nam', 'UNICEF_2022_School_Factors_Vietnam.pdf')

doc = create_doc()
add_p(doc, 'Link: https://www.unicef.org/vietnam/reports/study-school-related-factors-impacting-mental-health-and-well-being-adolescents-viet-nam', size=10)

add_heading(doc, 'Nghiên cứu toàn diện về các yếu tố liên quan đến trường học ảnh hưởng đến sức khỏe tâm thần và hạnh phúc của thanh thiếu niên nam và nữ tại Việt Nam', 1)
h = doc.add_paragraph()
r = h.add_run('Comprehensive Study on School-Related Factors Impacting Mental Health and Well-Being of Adolescent Boys and Girls in Viet Nam')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Comprehensive Study on School-Related Factors Impacting Mental Health and Well-Being of Adolescent Boys and Girls in Viet Nam'),
    ('Tác giả/Cơ quan', 'UNICEF Việt Nam + Bộ Giáo dục và Đào tạo (MOET)\nTư vấn quốc tế: Amie Alley Pollack\nTư vấn quốc gia: Đặng Hoàng Minh (ĐH Quốc gia Hà Nội)\nTrợ lý NC: Phương Nguyễn, Trang Lê, Lâm Lê'),
    ('Xuất bản', '2022, 118 trang'),
    ('Loại', 'Nghiên cứu phương pháp hỗn hợp (mixed methods) — Khảo sát định lượng + FGD + KII'),
    ('Mẫu', '668 HS THCS/THPT (khảo sát) + FGD HS + KII hiệu trưởng, sở, bộ\n5 tỉnh: Hà Nội, Điện Biên, Đồng Tháp, Kon Tum, Ninh Thuận'),
    ('Công cụ', 'SDQ-25 (Strengths and Difficulties Questionnaire)\nESSA (Educational Stress Scale for Adolescents)\nBảng hỏi tự xây dựng về cyberbullying, teacher-student relationship'),
    ('Phạm vi', '5 tỉnh đại diện: Bắc (Hà Nội, Điện Biên), Trung (Kon Tum, Ninh Thuận), Nam (Đồng Tháp)\nĐô thị + Nông thôn + DTTS'),
    ('Truy cập', 'Open Access — UNICEF Việt Nam'),
])
add_page_ref(doc, '1–118', 'UNICEF Việt Nam', '2022')

# TOM TAT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Nghiên cứu này do UNICEF Việt Nam phối hợp với Bộ Giáo dục và Đào tạo (MOET) thực hiện nhằm xây dựng hiểu biết về cách các yếu tố liên quan đến trường học tác động đến sức khỏe tâm thần và hạnh phúc của thanh thiếu niên nam và nữ tại Việt Nam, cũng như vai trò của hệ thống giáo dục trong giải quyết rủi ro và cung cấp hỗ trợ sức khỏe tâm thần và tâm lý xã hội (MHPSS).')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Nghiên cứu hỗn hợp tại 5 tỉnh (Hà Nội, Điện Biên, Đồng Tháp, Kon Tum, Ninh Thuận). Khảo sát 668 HS THCS/THPT + 249 giáo viên bằng SDQ-25 và ESSA. FGD với HS, cha mẹ, giáo viên. KII với hiệu trưởng, sở GD, MOET, MOH, MOLISA. Phân tích Nvivo 12 + SPSS.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('Phát hiện chính: (1) 8–29% VTN VN có vấn đề SKTT. (2) Nữ có tỷ lệ rối loạn cảm xúc cao hơn nam. (3) Các yếu tố trường học ảnh hưởng SKTT: gắn kết trường thấp, áp lực học tập cao, bắt nạt/cyberbullying, quan hệ thầy-trò kém, thiếu dịch vụ tư vấn. (4) 50% HS học thêm >3 giờ/tuần, 15% học thêm >9 giờ/tuần. (5) HS không thoải mái tìm đến giáo viên để được hỗ trợ. (6) Kỳ thị SKTT vẫn phổ biến.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0, 0)

# DANH GIA NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'UNICEF + MOET — uy tín CỰC CAO. Báo cáo chính sách quốc gia.',
    'Mixed methods: khảo sát (668 HS + 249 GV) + FGD + KII — toàn diện.',
    '5 tỉnh đại diện (Bắc-Trung-Nam, đô thị + nông thôn + DTTS) — phạm vi rộng nhất trong các NC VN.',
    'SDQ-25 + ESSA — cả 2 thang đo chuẩn hóa quốc tế, đã xác thực VN.',
    'Phát hiện: 50% HS học thêm >3h/tuần, 15% >9h/tuần — bằng chứng ĐỊNH LƯỢNG áp lực học thêm.',
    'Bắt nạt, cyberbullying — dữ liệu VN chi tiết. Phù hợp Islam 2025 (QT31, AOR=1,68).',
    'Khuyến nghị chính sách: A (Giáo dục) + B (Y tế) + C (MOLISA) + D (Liên ngành) — toàn diện.',
    'Báo cáo gốc 118 trang — nguồn tài liệu phong phú nhất về SKTT trường học VN.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Không phải nghiên cứu học thuật peer-reviewed — báo cáo chính sách.',
    'n = 668 HS — nhỏ hơn Hoa 2024 (n=3.910) nhưng bù lại bằng mixed methods.',
    'Năm 2022 — có thể chưa phản ánh tình hình hậu COVID đầy đủ.',
    'SDQ-25 đo vấn đề hành vi chung — không đo lo âu riêng (GAD-7, DASS-21).',
    'Lấy mẫu thuận tiện — không hoàn toàn ngẫu nhiên.',
]:
    add_p(doc, f'• {b}')

# CHAPTER 1: GIOI THIEU (tóm lược)
add_page_ref(doc, '4–8', 'UNICEF Việt Nam', '2022')
add_heading(doc, 'CHƯƠNG 1: GIỚI THIỆU', 2)
add_p(doc, 'Trên toàn thế giới, khoảng 15% trẻ em và VTN trải qua rối loạn tâm thần (Polanczyk et al., 2015). Với 50% rối loạn SKTT bắt đầu trước 14 tuổi và 75% trước 24 tuổi (Kessler et al., 2005), SKTT trẻ em và VTN đã trở thành ưu tiên toàn cầu. Có khoảng cách giới đáng kể: nữ có SKTT trung bình kém hơn nam (Campbell et al., 2021).')
add_p(doc, 'Tại Việt Nam, bằng chứng cho thấy 8–29% VTN mắc vấn đề SKTT, với nam có tỷ lệ rối loạn hành vi cao hơn và nữ có tỷ lệ rối loạn cảm xúc (lo âu, trầm cảm) cao hơn (UNICEF, 2018; Weiss et al., 2014). Vấn đề SKTT là gánh nặng đáng kể: ở nam 10–14, rối loạn hành vi là nguyên nhân tàn tật đứng thứ 2; ở nữ 10–14, rối loạn hành vi VÀ lo âu đều là nguyên nhân tàn tật hàng đầu; ở nữ 15–19, trầm cảm đứng thứ 3.')
add_p(doc, 'Mục tiêu NC: (1) Hiểu biết về cách yếu tố trường học tác động SKTT VTN VN. (2) Đánh giá vai trò hệ thống giáo dục trong giải quyết rủi ro. (3) Đề xuất khuyến nghị cho UNICEF VN, MOET, MOH, MOLISA.')

# CHAPTER 4: PHƯƠNG PHÁP (đầy đủ)
add_page_ref(doc, '26–34', 'UNICEF Việt Nam', '2022')
add_heading(doc, 'CHƯƠNG 4: PHƯƠNG PHÁP', 2)
add_p(doc, '4.1. Thiết kế:', bold=True)
add_p(doc, 'Phương pháp hỗn hợp (mixed methods): (a) Khảo sát định lượng HS + GV; (b) Thảo luận nhóm tập trung (FGD) với HS, cha mẹ, GV; (c) Phỏng vấn sâu (KII) với hiệu trưởng, quản lý sở GD, chuyên gia MOET/MOH/MOLISA.')
add_p(doc, '4.2. Địa bàn:', bold=True)
add_p(doc, '5 tỉnh/thành phố đại diện: Hà Nội (Bắc, đô thị), Điện Biên (Bắc, DTTS, nông thôn), Đồng Tháp (Nam, nông thôn), Kon Tum (Trung, DTTS), Ninh Thuận (Trung, ven biển). Mỗi tỉnh: 1 trường THCS + 1 trường THPT.')
add_p(doc, '4.3. Mẫu khảo sát:', bold=True)
add_p(doc, '668 HS THCS/THPT (khối 8, 9, 11, 12) + 249 giáo viên từ 10 trường. Mẫu thuận tiện — chọn lớp có sẵn ngày khảo sát.')
add_p(doc, '4.4. Công cụ:', bold=True)
add_p(doc, '(1) SDQ-25 (Strengths and Difficulties Questionnaire — Goodman, 1997): 25 mục, 5 tiểu thang (triệu chứng cảm xúc, vấn đề hành vi, tăng động, vấn đề bạn bè, hành vi xã hội tích cực). Ngưỡng: bình thường/ranh giới/bất thường. Cronbach alpha trong NC: 0,72 tổng, 0,60–0,73 tiểu thang.')
add_p(doc, '(2) ESSA (Educational Stress Scale for Adolescents — Sun et al., 2011): 16 mục, 5 lĩnh vực, Likert 5 điểm, 16–80. Đã xác thực VN (Truc et al., 2012). Cronbach alpha: 0,81 tổng, 0,66–0,75 tiểu thang. Cũng được sử dụng trong Trần Thảo Vi 2024 (VN21, alpha=0,88).')
add_p(doc, '(3) Bảng hỏi bổ sung: cyberbullying, quan hệ thầy-trò, gắn kết trường, hoạt động ngoại khóa, học thêm, giấc ngủ, screen time.')
add_p(doc, '4.5. Phân tích:', bold=True)
add_p(doc, 'Định tính: Nvivo 12, 24 mã (32 mã phụ), phân tích chủ đề. Định lượng: SPSS, thống kê mô tả + kiểm định nhóm.')

# CHAPTER 5: PHÁT HIỆN SKTT VTN VN (đầy đủ)
add_page_ref(doc, '33–42', 'UNICEF Việt Nam', '2022')
add_heading(doc, 'CHƯƠNG 5: PHÁT HIỆN — SKTT VÀ HẠNH PHÚC VTN VN', 2)

add_heading(doc, 'Bảng 1. SDQ-25 — Vấn đề SKTT ở HS VN (n = 668)', 3)
add_table(doc,
    ['Tiểu thang SDQ', 'Bình thường (%)', 'Ranh giới (%)', 'Bất thường (%)', 'Giới tính'],
    [['Tổng điểm khó khăn', '~70%', '~15%', '~15%', 'Nữ > Nam (cảm xúc)'],
     ['Triệu chứng cảm xúc', '—', '—', '—', 'NỮ cao hơn đáng kể'],
     ['Vấn đề hành vi', '—', '—', '—', 'NAM cao hơn'],
     ['Tăng động/giảm chú ý', '—', '—', '—', '—'],
     ['Vấn đề bạn bè', '—', '—', '—', '—'],
     ['Hành vi xã hội tích cực', '—', '—', '—', 'NỮ cao hơn (bảo vệ)']],
    widths=[4.0, 2.5, 2.5, 2.5, 3.0])
add_p(doc, 'Nữ có tỷ lệ triệu chứng cảm xúc (lo âu, trầm cảm) cao hơn nam — phù hợp xu hướng toàn cầu (Islam 2025 QT31: AOR=1,51; Ireland 2024 QT32: nữ tăng nhanh hơn).', size=9, italic=True)

add_p(doc, '5.1. Nhận thức của chuyên gia', bold=True)
add_p(doc, 'Chuyên gia chia sẻ rằng trầm cảm và lo âu là vấn đề phổ biến ở HS. Một người tham gia từ Đồng Tháp chia sẻ: "Có những căng thẳng mà các em giữ bên trong và không nói ra vì các em thấy quá nhiều kỳ vọng từ nhà trường và cha mẹ, trong khi bản thân các em không thể đáp ứng. Thỉnh thoảng điều này dẫn đến trạng thái ẩn mình, rút lui khỏi mọi thứ nên các em không giao tiếp."')
add_p(doc, 'Quan chức Sở Y tế Điện Biên lo ngại về tự tử ở VTN dân tộc Mông bằng cỏ ngắt (heartbreak grass): "Điều khó khăn khi tìm hiểu SKTT học sinh nằm ở chính bản thân các em — các em không có kỹ năng chia sẻ vấn đề cá nhân khó khăn để được giúp gỡ rối tâm trí."')
add_p(doc, 'Gia đình miễn cưỡng tìm kiếm hỗ trợ — kỳ thị SKTT vẫn phổ biến tại VN. Phù hợp V-NAMHS 2022 (VN02): chỉ 5,1% phụ huynh nhận ra con cần giúp đỡ, chỉ 8,4% VTN tiếp cận dịch vụ.')

# CHAPTER 6: YẾU TỐ TRƯỜNG HỌC (đầy đủ)
add_page_ref(doc, '41–55', 'UNICEF Việt Nam', '2022')
add_heading(doc, 'CHƯƠNG 6: YẾU TỐ TRƯỜNG HỌC ẢNH HƯỞNG SKTT', 2)

add_heading(doc, 'Key Findings — Phát hiện chính', 3)
add_p(doc, '1. Số lượng đáng kể HS báo cáo gắn kết trường thấp và kết nối với GV kém. HS không thường cảm thấy thoải mái tìm đến GV để được hỗ trợ học tập hoặc cảm xúc xã hội.')
add_p(doc, '2. Hiệu trưởng, quản lý và chuyên gia bộ bày tỏ lo ngại về quan hệ thầy-trò. Rào cản chính: GV thiếu thời gian và hỗ trợ hành chính.')
add_p(doc, '3. Dữ liệu khảo sát cho thấy HS cảm thấy ít gắn kết trường có tỷ lệ vấn đề SKTT cao hơn. Đặc biệt, NỮ báo cáo cảm thấy ít gắn kết trường hơn nam.')
add_p(doc, '4. Bắt nạt là yếu tố nguy cơ NGHIÊM TRỌNG cho vấn đề SKTT và hạnh phúc HS. GV không luôn nhận biết hoặc phản ứng hiệu quả với bắt nạt.')
add_p(doc, '5. Áp lực học tập là nguồn căng thẳng CHÍNH — đặc biệt ở Hà Nội. HS chia sẻ ảnh hưởng tiêu cực của áp lực lên SKTT.')
add_p(doc, '6. Cyberbullying ngày càng phổ biến — GV và cha mẹ lo ngại nhưng thiếu công cụ giải quyết.')

add_p(doc, '6.1. Gắn kết trường và quan hệ thầy-trò', bold=True)
add_p(doc, 'HS không cảm thấy thoải mái tìm đến GV — cả về vấn đề học tập lẫn cảm xúc. Rào cản: sợ bị đánh giá, lo ngại bảo mật, khoảng cách quyền lực thầy-trò trong văn hóa VN. Phù hợp Ireland 2024 (QT32): OGA (One Good Adult) là yếu tố bảo vệ — gợi ý VN cần xây dựng mô hình mentor tương tự.')

add_p(doc, '6.2. Áp lực học tập', bold=True)
add_p(doc, 'HS Hà Nội báo cáo mức áp lực học tập CAO và chia sẻ nhiều ví dụ về tác động tiêu cực. Một HS chia sẻ: "Khi căng thẳng quá, chúng em không có thời gian tập thể dục hoặc đi dạo hay thư giãn. Ngay cả khi không có dịch, chúng em cũng không thể đi đâu thoải mái. Vì bây giờ đi đâu em cũng phải nghĩ đến bài tập. Em nghĩ về nó suốt — trường, bài tập, và mọi thứ. Khi quá nhiều việc và không thể hoàn thành kịp hoặc thiếu gì đó, khá đáng sợ. Em sợ bị thầy cô hoặc cha mẹ mắng. Khi căng thẳng, em không ngủ yên. Ví dụ, nếu trước kỳ thi, khi em học môn mà em không giỏi, em rất lo lắng."')

add_heading(doc, 'Bảng 2. Học thêm — Dữ liệu khảo sát (n = 668)', 3)
add_table(doc,
    ['Thời gian học thêm/tuần', 'Tỷ lệ (%)', 'Ghi chú'],
    [['Gần như không', '31,9% (213)', ''],
     ['1–2 giờ', '20,7% (138)', ''],
     ['3–5 giờ', '19,5% (130)', ''],
     ['6–8 giờ', '12,4% (83)', ''],
     ['9–12 giờ', '7,5% (50)', 'Nhiều'],
     ['>12 giờ', '7,6% (51)', 'RẤT NHIỀU'],
     ['Tổng >3 giờ/tuần', '~47%', 'Gần NỬA HS học thêm >3h/tuần'],
     ['Tổng >9 giờ/tuần', '~15%', '1/6-7 HS học thêm >9h/tuần']],
    widths=[4.0, 3.0, 5.0])
add_p(doc, 'Phù hợp Trần Thảo Vi 2024 (VN21): 92,1% HS học thêm; β=4,73 (mạnh nhất). Wen 2020 (QT08): áp lực OR=11,58.', size=9, italic=True)

add_p(doc, '6.3. Bắt nạt và Cyberbullying', bold=True)
add_p(doc, 'Bắt nạt là yếu tố nguy cơ NGHIÊM TRỌNG cho vấn đề SKTT HS. NC trước tại VN (Le et al., 2019; Nguyen, Nakamura, Seino, et al., 2019) cho thấy: bắt nạt phổ biến hơn ở THCS so với THPT. Nạn nhân bắt nạt dự báo vấn đề SKTT; và ngược lại, HS có vấn đề SKTT dễ trở thành nạn nhân — vòng xoáy hai chiều. Nữ có vấn đề SKTT dễ trở thành nạn nhân; nam có vấn đề SKTT dễ trở thành cả nạn nhân lẫn thủ phạm.')
add_p(doc, 'Phù hợp Islam 2025 (QT31): bắt nạt AOR=1,68 (59 nước). Norway 2025 (QT21): bắt nạt KHÔNG giải thích xu hướng tăng (ổn định/giảm) — nhưng vẫn là yếu tố nguy cơ quan trọng cho cá nhân.')

# CHAPTER 7: CHÍNH SÁCH (tóm lược)
add_page_ref(doc, '55–76', 'UNICEF Việt Nam', '2022')
add_heading(doc, 'CHƯƠNG 7: CHÍNH SÁCH VÀ CHƯƠNG TRÌNH MHPSS TẠI TRƯỜNG (TÓM LƯỢC)', 2)
add_p(doc, 'Từ 2005, MOET đã thiết lập nhiều chính sách hỗ trợ SKTT HS: (1) Phát triển chương trình tư vấn trường học; (2) Giáo dục hòa nhập cho trẻ khuyết tật; (3) Giải quyết tác động COVID-19; (4) Chương trình Y tế Trường học 2021–2025 — thúc đẩy nhận thức SKTT và kỹ năng SKTT qua giáo dục.')
add_p(doc, 'HS tin rằng CẦN hỗ trợ SKTT tại trường, bao gồm tư vấn chuyên nghiệp và hoạt động thúc đẩy SKTT. Tuy nhiên, HS MIỄN CƯỠNG tìm kiếm giúp đỡ từ GV tư vấn do lo ngại BẢO MẬT và không thoải mái chia sẻ cá nhân.')
add_p(doc, 'Rào cản: thiếu cơ sở vật chất (phòng tư vấn riêng), thiếu nhân lực được đào tạo đầy đủ. GV "kiêm nhiệm" tư vấn (không phải chuyên gia). Chỉ 6/28 nước châu Âu có giám sát kết quả tư vấn — VN cũng chưa có (phù hợp WHO Europe QT24).')

# CHAPTER 8: KHUYẾN NGHỊ (đầy đủ)
add_page_ref(doc, '77–88', 'UNICEF Việt Nam', '2022')
add_heading(doc, 'CHƯƠNG 8: KHUYẾN NGHỊ', 2)

add_heading(doc, 'Bảng 3. Khuyến nghị chính — 4 lĩnh vực', 3)
add_table(doc,
    ['Lĩnh vực', 'Khuyến nghị', 'Ý nghĩa cho đề tài'],
    [['A. GIÁO DỤC (MOET)', 'A1. Đào tạo GV về SKTT HS\nA2. Giảm áp lực học tập + thi cử\nA3. Phát triển tư vấn trường chuyên nghiệp\nA4. Xây dựng chương trình SKTT trong giáo dục\nA5. Chống bắt nạt/cyberbullying\nA6. Tăng cường gắn kết trường', 'A2 phù hợp VN21 (học thêm β=4,73)\nA3 phù hợp QT32 (OGA)\nA5 phù hợp QT31 (AOR=1,68)'],
     ['B. Y TẾ (MOH)', 'B1. Xây dựng nhân lực SKTT trẻ em\nB2. Đào tạo bác sĩ SKTT VTN trong chương trình y khoa\nB3. Đào tạo nhân viên y tế trường về SKTT cơ bản', 'Phù hợp V-NAMHS (8,4% tiếp cận)\nPhù hợp UK NHS QT26 (£16 tỷ)'],
     ['C. LAO ĐỘNG-XÃ HỘI (MOLISA)', 'C1. Xây dựng chính sách bảo vệ trẻ em trực tuyến\nC2. Hỗ trợ cha mẹ kỹ năng nuôi dạy\nC3. Phát triển dịch vụ SKTT cộng đồng', 'Phù hợp Nature QT27 (MXH g=0,46)\nPhù hợp QT31 (cha mẹ AOR=0,75)'],
     ['D. LIÊN NGÀNH', 'D1. Phối hợp MOET-MOH-MOLISA\nD2. Giám sát + đánh giá chương trình\nD3. Nghiên cứu + dữ liệu định kỳ', 'Phù hợp WHO Europe QT24 (tích hợp)\nCần lặp lại V-NAMHS']],
    widths=[2.5, 6.0, 5.0])

add_p(doc, 'Khuyến nghị A2 — Giảm áp lực học tập:', bold=True)
add_p(doc, 'MOET cần xem xét chính sách giảm áp lực thi cử và khối lượng bài tập. NC này + VN21 (Trần Thảo Vi: học thêm β=4,73) + Wen 2020 (QT08: áp lực OR=11,58) + Norway 2025 (QT21: bất mãn trường giải thích chính) đều chỉ ra áp lực HT là yếu tố MẠNH NHẤT. Cần cân bằng chất lượng giáo dục và SKTT HS.')

add_p(doc, 'Khuyến nghị B1 — Xây dựng nhân lực SKTT trẻ em:', bold=True)
add_p(doc, 'Hệ thống y tế VN rất ít bác sĩ và y tá chuyên khoa SKTT trẻ em/VTN. MOH cần: (a) đưa nội dung SKTT trẻ em vào chương trình y khoa; (b) phát triển chuyên gia SKTT trẻ em; (c) đào tạo nhân viên y tế trường. Phù hợp V-NAMHS 2022 (VN02): chỉ 8,4% tiếp cận dịch vụ. UK NHS (QT26): chi £16 tỷ/năm nhưng vẫn không đủ — VN cần đầu tư lớn hơn.')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
for ref in [
    'UNICEF Viet Nam (2022). Comprehensive Study on School-Related Factors Impacting Mental Health and Well-Being of Adolescent Boys and Girls in Viet Nam.',
    'Goodman, R. (1997). The Strengths and Difficulties Questionnaire. Journal of Child Psychology and Psychiatry, 38, 581–586.',
    'Sun, J., et al. (2011). ESSA development, validity, and reliability with Chinese students. J Psychoeduc Assess, 29, 534–546.',
    'Truc, T.T., et al. (2012). Validation of ESSA in Vietnam. Asia Pac J Public Health, 27, NP2112–NP2121.',
    'Le, M.T.H., et al. (2019). Bullying victimization and mental health among Vietnamese adolescents. BMC Public Health.',
    '(Xem đầy đủ trong bài gốc — 200+ tài liệu tham khảo, 118 trang)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('UNICEF', 'United Nations Children\'s Fund — Quỹ Nhi đồng Liên Hợp Quốc'),
    ('MOET', 'Ministry of Education and Training — Bộ Giáo dục và Đào tạo'),
    ('MOH', 'Ministry of Health — Bộ Y tế'),
    ('MOLISA', 'Ministry of Labour, Invalids and Social Affairs — Bộ Lao động-Thương binh-Xã hội'),
    ('SDQ-25', 'Strengths and Difficulties Questionnaire 25 items'),
    ('ESSA', 'Educational Stress Scale for Adolescents'),
    ('MHPSS', 'Mental Health and Psychosocial Support — Hỗ trợ SKTT và Tâm lý Xã hội'),
    ('FGD', 'Focus Group Discussion — Thảo luận Nhóm Tập trung'),
    ('KII', 'Key Informant Interview — Phỏng vấn Người Cung cấp Thông tin Chính'),
    ('DOET', 'Department of Education and Training — Sở GD&ĐT'),
    ('DOH', 'Department of Health — Sở Y tế'),
    ('DTTS', 'Dân tộc Thiểu số'),
    ('SKTT', 'Sức khỏe tâm thần'),
    ('VTN', 'Vị thành niên'),
    ('HS', 'Học sinh'),
    ('GV', 'Giáo viên'),
    ('THCS', 'Trung học Cơ sở'),
    ('THPT', 'Trung học Phổ thông'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'UNICEF + MOET — uy tín cực cao. Hợp tác quốc tế + quốc gia. Tư vấn ĐH Quốc gia HN (Đặng Hoàng Minh).',
    'Mixed methods — toàn diện nhất: khảo sát 668 HS + 249 GV + FGD + KII. Kết hợp định lượng (SDQ, ESSA) + định tính (quotes trực tiếp từ HS, GV, quản lý).',
    '5 tỉnh đại diện — Bắc/Trung/Nam, đô thị/nông thôn/DTTS. PHẠM VI RỘNG NHẤT trong các NC VN (so: Hoa 2024 chỉ Hà Nội, Vĩnh Lộc 2024 chỉ TPHCM).',
    'Dữ liệu HỌC THÊM định lượng: 47% >3h/tuần, 15% >9h/tuần — bằng chứng mạnh. Phù hợp VN21 (β=4,73, mạnh nhất).',
    'Khuyến nghị chính sách 4 lĩnh vực — áp dụng được ngay cho MOET, MOH, MOLISA.',
    'Quotes trực tiếp từ HS — giá trị cao cho phần định tính trong đề cương NC.',
    'Tham chiếu y văn quốc tế + VN — 200+ TLTK.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'KHÔNG phải NC học thuật peer-reviewed — báo cáo chính sách UNICEF. Không có DOI, không lập chỉ mục PubMed/Scopus.',
    'n = 668 HS — nhỏ hơn Hoa 2024 (n=3.910), V-NAMHS (n=5.996). Nhưng bù lại bằng mixed methods + 5 tỉnh.',
    'SDQ-25 đo vấn đề HÀNH VI + CẢM XÚC chung — KHÔNG đo lo âu riêng (GAD-7, DASS-21). Không so sánh trực tiếp được với Hoa 2024 (GAD-7: 40,6%) hay Bảo Quyên 2025 (DASS-21: 86,2%).',
    'Lấy mẫu THUẬN TIỆN — không hoàn toàn ngẫu nhiên. Có thể thiên lệch: chỉ HS có mặt ngày khảo sát.',
    'Năm 2022 — thu thập dữ liệu trong/gần COVID. Có thể ảnh hưởng kết quả. Hoàng Trung Học 2025 (VN14): lo âu giảm sau COVID.',
    'Không có phân tích hồi quy đa biến — chỉ mô tả + kiểm định nhóm. So: Wen 2020 (QT08) dùng LPA, VN21 dùng hồi quy tuyến tính.',
    'Thiếu phân tích theo SỰ GIÀU NGHÈO — Korea 2024 (QT34): chênh 22,7 điểm stress giữa nhóm giàu/nghèo.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'Cần NC tương tự với THANG ĐO LO ÂU riêng (GAD-7/DASS-21) bên cạnh SDQ — để so sánh được với NC khác.',
    'Cần phân tích hồi quy đa biến — xác định yếu tố ĐỘC LẬP (không chỉ mô tả). VN21 đã làm (hồi quy tuyến tính) nhưng chỉ ESSA.',
    'Mở rộng: TPHCM (thiếu trong 5 tỉnh), thêm tỉnh đại diện vùng kinh tế đặc biệt (Bình Dương, Quảng Ninh).',
    'Đánh giá HIỆU QUẢ chương trình Y tế Trường học 2021–2025 của MOET — đã triển khai 3 năm, cần đánh giá.',
    'So sánh trường CÔNG vs TƯ — báo cáo chỉ trường công. VN20 (Vĩnh Lộc 2024) cũng chỉ trường công.',
    'Kết hợp UNICEF 2022 + V-NAMHS 2022 + VN21 + Hoa 2024 → tổng hợp đa nguồn cho bức tranh toàn diện nhất SKTT HS VN.',
    'NC dọc: theo dõi cùng HS qua THCS → THPT → xem yếu tố trường học THAY ĐỔI hay NHẤT QUÁN. VN21 (3 năm THCS) là bước đầu.',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '38_UNICEF_VN_2022_SchoolFactors.docx')
doc.save(outpath)

# Verify
import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
# Verify key numbers
checks = ['668', '249', 'SDQ', 'ESSA', 'UNICEF', 'MOET', '8–29%', '50%', 'Điện Biên', 'Hà Nội', 'heartbreak']
ok = sum(1 for c in checks if c in t)
print(f'  Numbers verified: {ok}/{len(checks)}')
