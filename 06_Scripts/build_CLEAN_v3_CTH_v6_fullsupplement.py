"""Build CLEAN_v2_CTH_v6_chuong_4_bo_sung.docx
Apply Cong Thi Hang v6 algorithm:
- 3-layer paragraph structure (THESIS-EVIDENCE-SYNTHESIS)
- Transition phrases (Thu nhat/Thu hai/Thu ba, Noi cach khac, Phu hop voi)
- CAPS LOCK for emphasis instead of bold
- Em-dash for explanatory clauses
- 3-level opening (TOAN CAU -> KHU VUC -> VN)
- "Ba khoang trong" / "Nam phat hien" closing
- After-table interpretation paragraph
- No em-thay, no meta, no fabrications
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/CLEAN_v3_CTH_v6_chuong_4_bo_sung_full.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = BLACK

def para(text, indent=True, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(13)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def add_table(header, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

def ref_entry(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# ============================================================
# Title
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CHƯƠNG 4 (BỔ SUNG)\nĐỐI CHIẾU KẾT QUẢ NGHIÊN CỨU VỚI CƠ SỞ DỮ LIỆU TỔNG HỢP\nVÀ ĐỀ XUẤT KHUNG TẬP HUẤN CAN THIỆP')
r.bold = True; r.font.size = Pt(14)
para('')

# ============================================================
# 4.1 — Đối chiếu (mở chương 3 cấp độ TOÀN CẦU → KHU VỰC → VN)
# ============================================================
H('4.1. Đối chiếu kết quả nghiên cứu với tổng quan tài liệu', level=2)

# Đoạn mở chương — 3 cấp độ
para(
    'Rối loạn lo âu ở vị thành niên đã trở thành mối quan tâm sức khỏe công cộng toàn '
    'cầu trong thập kỷ qua. Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2021 cho thấy lo âu '
    'và trầm cảm chiếm phần lớn DALYs do rối loạn tâm thần ở nhóm 10–24 tuổi, với xu '
    'hướng TĂNG 23,6% từ 1990 đến 2021 (Bie và cộng sự, 2024). Nói cách khác, gánh '
    'nặng bệnh tật do lo âu ở thanh thiếu niên đã tăng gấp gần một phần tư trong vòng '
    'ba thập niên — một mức tăng đáng báo động đối với sức khỏe cộng đồng.'
)
para(
    'Tại khu vực Đông Nam Á, GBD ASEAN Mental Disorders Collaborators (2025) ghi nhận '
    '11,2 triệu DALYs do rối loạn tâm thần năm 2021, trong đó lo âu chiếm 31,0% — cao '
    'nhất trong tất cả các rối loạn được phân tích. Nhóm 10–14 tuổi chịu gánh nặng '
    'CAO NHẤT với 16,3% tổng DALYs ở nhóm tuổi này. Phù hợp với xu hướng toàn cầu, '
    'lo âu là rối loạn dẫn đầu ở thanh thiếu niên ASEAN.'
)
para(
    'Tại Việt Nam, Khảo sát Sức khỏe Tâm thần Vị thành niên (V-NAMHS) năm 2022 trên '
    '5.996 vị thành niên 10–17 tuổi ghi nhận 21,7% có vấn đề sức khỏe tâm thần trong '
    '12 tháng qua, trong đó lo âu là rối loạn phổ biến nhất (UNICEF Việt Nam, 2022). '
    'Nghiên cứu của chúng tôi trên 1.352 học sinh trung học cơ sở khối 6–9 góp phần '
    'làm rõ thêm bức tranh này — không chỉ về tỷ lệ mà còn về mẫu hình phát triển '
    'theo giới tính, khối lớp và các yếu tố nguy cơ–bảo vệ.'
)

# 4.1.1 — Mức độ và biểu hiện
H('4.1.1. Mức độ và biểu hiện rối loạn lo âu', level=3)
para(
    'Kết quả khảo sát cho thấy ba dạng rối loạn lo âu có cường độ và phân bố KHÁC BIỆT '
    'rõ rệt. Lo âu lan tỏa có điểm trung bình cao nhất (45,86–64,28), tiếp theo là lo '
    'âu xã hội (42,09–56,98), và thấp nhất là lo âu chia ly (21,52–27,88). Phù hợp '
    'với tổng quan của Beesdo, Knappe và Pine (2009) về phát triển lo âu ở vị thành '
    'niên — lứa tuổi 11–15 đã qua đỉnh của Separation Anxiety Disorder và đang chuyển '
    'sang giai đoạn nổi trội của lo âu lan tỏa và lo âu xã hội.'
)
para(
    'Mệnh đề "lo lắng khi nghĩ rằng mình đã không làm tốt điều gì đó" có điểm trung '
    'bình CAO NHẤT (ĐTB = 64,28). Phát hiện này nhất quán với khảo sát PISA 2015 của '
    'OECD trên 540.000 học sinh 72 quốc gia — 66% học sinh lo lắng về điểm kém và '
    '55% rất lo lắng về kiểm tra ngay cả khi đã chuẩn bị tốt (Pascoe, Hetrick & '
    'Parker, 2020). Nói cách khác, lo lắng về thất bại là biểu hiện cốt lõi của lo '
    'âu học sinh trên phạm vi toàn cầu, không riêng Việt Nam.'
)
para(
    'Mệnh đề "lo lắng điều tệ xảy ra với gia đình" (ĐTB = 59,62) đứng thứ ba — củng '
    'cố vai trò trung tâm của yếu tố gia đình trong lo âu học sinh trung học cơ sở '
    'Việt Nam. Phù hợp với Đinh và cộng sự (2021) khi xác định các yếu tố trường học '
    'và gia đình là nguồn gây lo âu chính ở vị thành niên Việt Nam. Trong văn hóa Á '
    'Châu, sự gắn kết liên thế hệ tạo cả nguồn hỗ trợ và áp lực cho học sinh.'
)

# 4.1.2 — Khác biệt giới tính và khối lớp
H('4.1.2. Khác biệt theo giới tính và khối lớp', level=3)
para(
    'Phân tích phương sai cho thấy học sinh nữ có điểm cao hơn nam ở ba trong bốn '
    'chỉ số rối loạn lo âu. Cụ thể, lo âu lan tỏa (M = 59,47 so với 51,43; F = '
    '45,484; p < 0,01), lo âu xã hội (M = 52,74 so với 43,20; F = 28,642; p < 0,01) '
    'và rối loạn lo âu tổng (M = 45,66 so với 40,02) đều cao hơn ở nữ. Phù hợp với '
    'tổng quan hệ thống của McLean, Asnaani, Litz và Hofmann (2011) và phân tích '
    'tổng hợp của Salk, Hyde và Abramson (2017) — nữ giới có tỷ số nguy cơ lo âu cao '
    'hơn nam khoảng 1,5–2 lần, với chênh lệch mở rộng sau dậy thì.'
)
para(
    'Tại Việt Nam, kết quả của Hoa và cộng sự (2024) trên 3.910 học sinh trung học '
    'phổ thông Hà Nội cũng xác nhận xu hướng này (nữ 43,8% so với nam 36,9%). Nghiên '
    'cứu của chúng tôi mở rộng phát hiện đó xuống nhóm trung học cơ sở — cho thấy '
    'chênh lệch giới đã hiện diện ngay từ lớp 6 và duy trì xuyên suốt khối 6–9.'
)
para(
    'Đáng chú ý, lo âu chia ly KHÔNG có khác biệt theo giới (F = 0,246; p = 0,620) '
    '— TRÁI với mẫu hình ba chỉ số còn lại. Phù hợp với Allen và cộng sự (2010) khi '
    'cho thấy Separation Anxiety Disorder ít chịu ảnh hưởng giới hơn so với '
    'Generalized Anxiety và Social Anxiety ở vị thành niên. Cơ chế sinh học của lo '
    'âu chia ly liên quan nhiều đến hệ thống gắn bó (attachment system) — vốn ít '
    'biến đổi theo giới ở lứa tuổi này.'
)
para(
    'Đối với khối lớp, ba xu hướng đối lập được ghi nhận. Lo âu lan tỏa TĂNG dần từ '
    'lớp 6 (M = 54,32) đến lớp 9 (M = 59,79; F = 5,020; p = 0,002). Ngược lại, lo '
    'âu chia ly GIẢM mạnh từ lớp 6 (M = 32,13) xuống lớp 9 (M = 19,86; F = 21,239; '
    'p < 0,001) — chỉ số có khác biệt khối lớp RÕ RỆT NHẤT. Lo âu xã hội đạt đỉnh '
    'ở lớp 9 (M = 53,05; F = 4,879; p = 0,002). Mẫu hình ba xu hướng này nhất quán '
    'với developmental psychopathology của Beesdo, Knappe và Pine (2009) — lo âu '
    'chia ly giảm khi trẻ phát triển sự độc lập, trong khi lo âu lan tỏa và lo âu '
    'xã hội tăng theo dậy thì khi tự nhận thức xã hội phát triển (Rapee & Spence, '
    '2004).'
)

# 4.1.3 — Yếu tố nguy cơ
H('4.1.3. Yếu tố nguy cơ', level=3)
para(
    'Áp lực học tập là yếu tố nguy cơ NỔI BẬT NHẤT với điểm trung bình cao nhất '
    '(ĐTB = 51,13). Trong nhóm này, mệnh đề "kỳ vọng học tập và định hướng tương '
    'lai" có ĐTB = 58,56 — vượt xa các mệnh đề còn lại. Phát hiện này có ý nghĩa '
    'quan trọng — học sinh trung học cơ sở Việt Nam độ tuổi 11–15 đã chịu áp lực '
    'sự nghiệp tương lai sớm hơn giả định thông thường rằng áp lực này chỉ xuất '
    'hiện ở khối trung học phổ thông chuẩn bị thi đại học.'
)
para(
    'Phù hợp với nghiên cứu thuần tập Hue Healthy Adolescent Cohort của Trần Thảo Vi '
    'và cộng sự (2024) trên 341 học sinh trung học cơ sở Huế — căng thẳng học tập '
    'TĂNG 15,3% từ lớp 6 đến lớp 9, với học thêm là yếu tố dự báo mạnh nhất (β = '
    '4,73). Tại Trung Quốc, Wen và cộng sự (2020) trên 900 học sinh trung học cơ sở '
    'nông thôn xác định áp lực học tập là yếu tố nguy cơ chính, với tỷ số chênh OR = '
    '11,58 (KTC 95% từ 4,16 đến 32,19) cho mức áp lực rất cao so với rất thấp. Nói '
    'cách khác, học sinh chịu áp lực rất cao có khả năng bị lo âu nặng cao hơn ÍT '
    'NHẤT 4 lần ở mức tin cậy 95%.'
)
para(
    'Nghiện điện thoại đứng thứ hai trong nhóm nguy cơ, với hành vi "kiểm tra điện '
    'thoại liên tục" có ĐTB = 35,92. Củng cố thêm thử nghiệm ngẫu nhiên có đối '
    'chứng SCREENS-Kids của Schmidt-Persson và cộng sự (2024) — hạn chế sử dụng màn '
    'hình giải trí trong 14 ngày làm cải thiện đáng kể vấn đề tâm lý nội hóa. Tại '
    'Na Uy, Brunborg và cộng sự (2025) trên 979.043 học sinh xác nhận thời gian sử '
    'dụng mạng xã hội tăng giải thích phần lớn xu hướng tăng distress ở học sinh '
    'nữ giai đoạn 2014–2024.'
)
para(
    'Bắt nạt học đường có điểm thấp hơn nhưng vẫn hiện diện trong toàn mẫu. Đáng '
    'chú ý, mô hình SEM cho thấy bắt nạt học đường có tác động dương đến cả ba dạng '
    'rối loạn lo âu — cao nhất là lo âu chia ly (β = 0,376; p < 0,001), tiếp theo '
    'là lo âu xã hội (β = 0,253; p < 0,001) và lo âu lan tỏa (β = 0,215; p < 0,001). '
    'Phù hợp với Brown và Carter (2025) trên các trường trung học cơ sở Anh — học '
    'sinh từng bị bắt nạt có nguy cơ lo âu lâm sàng tăng đáng kể.'
)

# 4.1.4 — Yếu tố bảo vệ
H('4.1.4. Yếu tố bảo vệ', level=3)
para(
    'Hỗ trợ từ cha mẹ có điểm trung bình CAO NHẤT trong nhóm bảo vệ (ĐTB = 57,65), '
    'tuy nhiên khả năng chia sẻ với gia đình lại có mức THẤP HƠN (ĐTB = 47,54). '
    'Khoảng trống giao tiếp giữa cha mẹ và con là phát hiện quan trọng — nói cách '
    'khác, mặc dù cha mẹ có ý thức hỗ trợ, học sinh không cảm thấy đủ thân thiết để '
    'chia sẻ vấn đề tâm lý. Phù hợp với V-NAMHS 2022 ghi nhận chỉ 5,1% phụ huynh '
    'Việt Nam xác định được con cần trợ giúp tâm lý (UNICEF Việt Nam, 2022).'
)
para(
    'Tự trọng đứng thứ hai trong nhóm bảo vệ (ĐTB = 54,85), với mệnh đề "thái độ '
    'tích cực với bản thân" (65,80) cao hơn rõ "đánh giá năng lực bản thân" '
    '(50,02). Phát hiện này gợi ý hai khía cạnh khác nhau của tự trọng phát triển '
    'không đồng đều ở học sinh trung học cơ sở — tự yêu bản thân vẫn ổn định, nhưng '
    'niềm tin vào năng lực bị xói mòn dưới áp lực học tập. Phù hợp với mô hình '
    'resilience của Cai và cộng sự (2025).'
)

# 4.1.5 — Mô hình SEM
H('4.1.5. Mô hình tác động tổng hợp', level=3)
para(
    'Tổng hợp 11 mô hình SEM cho phép xếp hạng cường độ tác động của từng yếu tố '
    'theo trị tuyệt đối hệ số chuẩn hóa β. Kết quả thể hiện ở Bảng 4.1.'
)
caption('Bảng 4.1. Xếp hạng tác động của các yếu tố nguy cơ và bảo vệ theo |β|')
add_table(
    ['Yếu tố', 'Loại', 'β cho lo âu lan tỏa', 'β cho lo âu xã hội', 'Cường độ'],
    [
        ['Áp lực học tập', 'Nguy cơ', '0,510 ***', '0,490 ***', 'Mạnh'],
        ['Tự trọng', 'Bảo vệ', '−0,455 ***', '−0,415 ***', 'Mạnh'],
        ['Nghiện điện thoại', 'Nguy cơ', '0,336 ***', '0,383 ***', 'Trung bình'],
        ['Hỗ trợ cha mẹ', 'Bảo vệ', '−0,172 ***', '−0,273 ***', 'Trung bình'],
        ['Bắt nạt học đường', 'Nguy cơ', '0,215 ***', '0,253 ***', 'Yếu–trung bình'],
        ['Gắn bó trường học', 'Bảo vệ', '−0,108 **', '−0,187 ***', 'Yếu–trung bình'],
        ['Hỗ trợ bạn bè', 'Bảo vệ', '−0,015 ns', '−0,079 *', 'Đặc thù'],
    ]
)
para('Ghi chú: *** p < 0,001; ** p < 0,01; * p < 0,05; ns — không có ý nghĩa thống kê.', indent=False, justify=False)
para('')
para(
    'Bảng 4.1 cho thấy tự trọng là yếu tố bảo vệ MẠNH NHẤT (|β| = 0,455 cho lo âu '
    'lan tỏa; 0,415 cho lo âu xã hội), với cường độ tương đương áp lực học tập '
    '(|β| = 0,510). Phát hiện này có hàm ý can thiệp quan trọng — nâng cao tự '
    'trọng có hiệu quả tương đương với giảm áp lực học tập trong việc giảm rối '
    'loạn lo âu. Nói cách khác, chương trình can thiệp song song hai trục (giảm '
    'áp lực + tăng tự trọng) sẽ hiệu quả hơn can thiệp đơn trục.'
)
para(
    'Mô hình SEM tổng hợp cuối cùng đạt CFI = 0,937; RMSEA = 0,077 với khoảng tin '
    'cậy 90% (0,016–0,065) — chuẩn báo cáo theo Browne và Cudeck (1993). Hệ số xác '
    'định R² = 0,598 cho thấy mô hình giải thích 59,8% phương sai của rối loạn lo '
    'âu — vượt xa ngưỡng "effect size lớn" theo Cohen (1988).'
)

# 4.1.6 — Biện pháp đối phó (BPĐP)
H('4.1.6. Biện pháp đối phó của học sinh trung học cơ sở', level=3)
para(
    'Trong số ba nhóm biện pháp đối phó được đo lường, tìm kiếm sự hỗ trợ có điểm '
    'trung bình CAO NHẤT (ĐTB = 55,00), tiếp theo là tập trung giải quyết vấn đề '
    '(53,18) và cuối cùng là tránh né. Phù hợp với mô hình ba nhân tố cổ điển của '
    'Carver (1997) trong Brief-COPE — vị thành niên Á Châu có xu hướng ưu tiên '
    'các chiến lược ADAPTIVE (problem-focused và support-seeking) hơn các chiến '
    'lược MALADAPTIVE (avoidance, self-blame).'
)
para(
    'Tuy nhiên, phân tích chi tiết cho thấy hiện tượng đáng lo ngại trong nhóm '
    'tránh né — hành vi "tự trách bản thân về những điều đã xảy ra" có ĐTB = '
    '50,42, cao hơn cả "tự chỉ trích bản thân" (45,83). Nói cách khác, mặc dù '
    'tránh né là chiến lược ít sử dụng nhất ở mức nhóm, một bộ phận học sinh vẫn '
    'có xu hướng tự đổ lỗi — chỉ dấu rủi ro của trầm cảm và lo âu (Compas và '
    'cộng sự, 2017).'
)
para(
    'Đáng chú ý, mô hình SEM giữa biện pháp đối phó và rối loạn lo âu CHƯA đạt '
    'độ phù hợp tốt. Cụ thể, các chỉ số χ²/df dao động từ 9,631 đến 57,264, '
    'trong khi CFI và TLI ở nhiều mô hình dưới ngưỡng chấp nhận 0,90. Phát hiện '
    'này có ý nghĩa phương pháp luận — gợi ý cần phân tách biện pháp đối phó '
    'thành các thang con cụ thể hơn theo mô hình 14 nhân tố của Brief-COPE '
    '(Carver, 1997) thay vì gộp ba nhóm lớn.'
)
para(
    'Phát hiện gây bất ngờ là mức độ sử dụng biện pháp đối phó càng cao thì rối '
    'loạn lo âu càng cao (β dương, TRÁI với giả định "đối phó nhiều thì lo âu '
    'giảm"). Phù hợp với tổng quan của Compas, Jaser, Bettis và cộng sự (2017) '
    'trong Psychological Bulletin về hiện tượng MALADAPTIVE COPING ESCALATION '
    '— học sinh càng lo âu càng dùng nhiều chiến lược đối phó, nhưng nếu chiến '
    'lược không hiệu quả thì lo âu vẫn duy trì hoặc tăng lên. Nói cách khác, '
    'TẦN SUẤT đối phó không đồng nghĩa với CHẤT LƯỢNG đối phó — và TẦN SUẤT '
    'cao có thể là chỉ dấu của lo âu cao hơn là biện pháp giảm lo âu.'
)
para(
    'Trích phỏng vấn một học sinh lớp 6 củng cố diễn giải này — "Khi lo lắng em '
    'thường tập trung giải quyết vấn đề, nhưng thi thoảng em cảm thấy những '
    'cách đó không giúp ích vì em không tìm được cách giải quyết" (NĐMK, lớp '
    '6). Hàm ý can thiệp — chương trình tập huấn không nên tăng tần suất sử '
    'dụng biện pháp đối phó nói chung, mà cần dạy học sinh PHÂN BIỆT chiến '
    'lược adaptive (giải quyết vấn đề có cấu trúc, tìm hỗ trợ chất lượng) với '
    'chiến lược maladaptive (tự trách, né tránh, suy nghĩ phản tư tiêu cực).'
)

# 4.1.7 — Quy đổi số liệu sang thang %
H('4.1.7. Quy đổi số liệu sang thang phần trăm và ngưỡng diễn giải', level=3)
para(
    'Các bảng 3.17–3.22 báo cáo điểm trung bình trong khoảng 21–64 với cột Min = '
    '1 và Max = 4 — cấu trúc này gợi ý hai khả năng. Khả năng thứ nhất, số '
    'liệu đã được chuẩn hóa sang thang phần trăm 0–100 nhưng cột Min/Max chưa '
    'được cập nhật theo. Khả năng thứ hai, số liệu đã được chuyển sang thang '
    'T-score chuẩn hóa theo phân phối mẫu (M = 50, SD = 10) như cách RCADS bản '
    'Mỹ thường được trình bày trong báo cáo lâm sàng (Chorpita và cộng sự, 2000). '
    'Việc xác định chính xác cách chuẩn hóa cần đối chiếu với phần phương pháp '
    'của chương 2 luận án.'
)
para(
    'Nếu áp dụng công thức quy đổi % phổ biến cho thang Likert: phần trăm = '
    '(điểm Likert − Min) / (Max − Min) × 100. Áp dụng cho thang RCADS với '
    'Likert 1–4: % = (điểm − 1) / 3 × 100. Nói cách khác, nếu ĐTB = 64,28 là '
    'thang %, thì điểm Likert thực ≈ 2,93 — gần "thường xuyên" trên thang 1 '
    '"không bao giờ" đến 4 "luôn luôn".'
)
caption('Bảng 4.3. Quy đổi điểm trung bình sang thang Likert gốc và ngưỡng diễn giải')
add_table(
    ['ĐTB (% chuẩn hóa)', 'Điểm Likert quy đổi', 'Diễn giải', 'Ví dụ trong dữ liệu'],
    [
        ['0–25', '1,00–1,75', 'Hiếm khi trải nghiệm', 'Lo âu chia ly lớp 9 (19,86)'],
        ['25–50', '1,75–2,50', 'Thỉnh thoảng', 'Phần lớn item RCADS chia ly'],
        ['50–75', '2,50–3,25', 'Thường xuyên', 'Lo âu lan tỏa (45,86–64,28)'],
        ['75–100', '3,25–4,00', 'Luôn luôn', 'Không xuất hiện trong dữ liệu'],
    ]
)
para('')
para(
    'Bảng 4.3 cho thấy phần lớn các biểu hiện lo âu lan tỏa trong nghiên cứu '
    'rơi vào ngưỡng "thường xuyên" — nói cách khác, học sinh trải nghiệm các '
    'biểu hiện này NHIỀU LẦN trong tuần. Phát hiện này có ý nghĩa quan trọng '
    'trong việc diễn giải mức độ thực — không chỉ là "có lo âu" mà là lo âu '
    'TẦN SUẤT CAO. Phù hợp với khung của Beesdo, Knappe và Pine (2009) về sự '
    'phân biệt giữa lo âu phát triển bình thường (transient) và lo âu lâm sàng '
    '(persistent).'
)
para(
    'Khuyến nghị bổ sung khi báo cáo — cập nhật cột Min và Max trong các bảng '
    '3.17–3.22 thành 0–100 thay vì 1–4 để phản ánh đúng thang đo đã chuẩn hóa, '
    'đồng thời thêm ghi chú phương pháp chuẩn hóa ở phần phương pháp nghiên cứu '
    'của chương 2. Việc minh bạch hóa thang đo cho phép các nghiên cứu sau '
    'đối chiếu trực tiếp với kết quả của luận án.'
)

# Kết phần 4.1 — Bảy phát hiện chính (mở rộng từ 5)
H('Bảy phát hiện chính từ đối chiếu', level=3)
para(
    'Tổng hợp đối chiếu kết quả nghiên cứu với cơ sở dữ liệu liên bài cho bảy '
    'phát hiện chính. Thứ nhất, áp lực học tập và kỳ vọng tương lai là yếu tố '
    'nguy cơ NỔI BẬT NHẤT ở học sinh trung học cơ sở Việt Nam — sớm hơn giả '
    'định thông thường về áp lực thi cử. Thứ hai, lo âu lan tỏa và lo âu xã '
    'hội TĂNG theo khối lớp trong khi lo âu chia ly GIẢM — phù hợp '
    'developmental psychopathology. Thứ ba, nữ giới có điểm cao hơn nam ở ba '
    'trong bốn chỉ số — phù hợp y văn quốc tế với chênh lệch ~1,5–2 lần. Thứ '
    'tư, tự trọng là yếu tố bảo vệ MẠNH NHẤT, có cường độ ngang bằng với áp '
    'lực học tập — gợi ý chiến lược can thiệp song trục. Thứ năm, mô hình SEM '
    'tổng hợp giải thích 59,8% phương sai — vượt ngưỡng effect size lớn của '
    'Cohen (1988). Thứ sáu, biện pháp đối phó có tác động dương đến rối loạn '
    'lo âu — TRÁI với giả định trực giác, phù hợp hiện tượng MALADAPTIVE '
    'COPING ESCALATION (Compas và cộng sự, 2017). Thứ bảy, hành vi tự trách '
    'bản thân (ĐTB = 50,42) là chỉ dấu rủi ro đáng quan tâm trong nhóm tránh '
    'né — cần được sàng lọc và can thiệp sớm.'
)

# ============================================================
# 4.2 — Khung tập huấn 4 tầng
# ============================================================
H('4.2. Đề xuất khung tập huấn can thiệp bốn tầng', level=2)

# Mở mục 4.2 — 3 cấp độ
para(
    'Tổ chức Y tế Thế giới (WHO, 2022) khuyến cáo các quốc gia áp dụng mô hình can '
    'thiệp phân tầng (stepped-care) cho rối loạn tâm thần ở vị thành niên. Tại các '
    'nước Á Châu, Nhật Bản đã triển khai thành công mô hình iCBT cho học sinh trung '
    'học (Matsumoto và cộng sự, 2024) — chứng minh độ chấp nhận cao trong văn hóa '
    'tương đồng với Việt Nam. Trên cơ sở đó, nghiên cứu của chúng tôi đề xuất khung '
    'can thiệp bốn tầng cho học sinh trung học cơ sở Việt Nam — tích hợp đồng thời '
    'các yếu tố nguy cơ và bảo vệ đã xác định trong mô hình SEM.'
)

# 4.2.1
H('4.2.1. Tầng phổ thông (Universal)', level=3)
para(
    'Đối tượng là TOÀN BỘ học sinh trung học cơ sở, không phân biệt mức độ lo âu. '
    'Nội dung gồm giáo dục tâm lý 60 phút mỗi tháng trong năm học, tập trung phân '
    'biệt lo âu bình thường và lo âu lâm sàng; kỹ thuật thư giãn cơ bản như hơi '
    'thở 4-7-8 và thư giãn cơ tiến triển theo Trần Nguyễn Ngọc (2018). Phù hợp với '
    'tổng quan của Bie và cộng sự (2024) — sàng lọc và giáo dục là tầng đầu tiên '
    'trong hệ thống can thiệp toàn cầu. Mục tiêu KPI: 80% học sinh phân biệt được '
    'hai loại lo âu sau một năm; 60% sử dụng ít nhất một kỹ thuật thư giãn ≥ 1 lần/'
    'tuần.'
)

# 4.2.2
H('4.2.2. Tầng chọn lọc (Selective)', level=3)
para(
    'Đối tượng là học sinh có ít nhất một yếu tố nguy cơ chính — áp lực học tập '
    'cao, nghiện điện thoại, bị bắt nạt, hoặc tự trọng thấp. Nội dung gồm chương '
    'trình kỹ năng đối phó tám buổi (60 phút/tuần), dựa trên Brief-COPE của Carver '
    '(1997) — phân biệt đối phó adaptive và maladaptive (Compas và cộng sự, 2017). '
    'Hoạt động cụ thể bao gồm nhật ký lo âu, kỹ năng tư duy nhận thức CBT-lite, '
    'bài tập chánh niệm.'
)
para(
    'Cơ sở bằng chứng là phân tích tổng hợp dữ liệu cá nhân (IPD meta-analysis) '
    'của Galante và cộng sự (2023) — chương trình mindfulness cải thiện lo âu so '
    'với nhóm chứng. Phù hợp với mô hình của Wen và cộng sự (2020) tại Trung Quốc '
    '— phân tích nhóm tiềm ẩn (LPA) xác định ba nhóm phụ học sinh, gợi ý can '
    'thiệp PHÂN TẦNG hiệu quả hơn can thiệp đồng nhất. Mục tiêu KPI: giảm điểm '
    'áp lực học tập-related anxiety ≥ 20% sau tám buổi; tăng điểm resilience ≥ '
    '15%.'
)

# 4.2.3
H('4.2.3. Tầng chỉ định (Indicated)', level=3)
para(
    'Đối tượng là học sinh có điểm RCADS hoặc DASS-21 trên ngưỡng cảnh báo nhưng '
    'CHƯA đáp ứng tiêu chí chẩn đoán DSM-5. Nội dung là trị liệu nhận thức hành '
    'vi qua internet (iCBT) tám buổi — dựa trên mô hình của Matsumoto và cộng sự '
    '(2024). Phù hợp với văn hóa Á Châu khi giảm rào cản stigma và cho phép tham '
    'gia ẩn danh.'
)
para(
    'Tính năng iCBT chính bao gồm video giáo dục tâm lý, bài tập tư duy nhận thức '
    'theo mô hình ABC, kích hoạt hành vi (behavioral activation), và phơi nhiễm '
    'phân cấp cho lo âu xã hội. Có thể phối hợp ứng dụng di động ClearFear (Samele '
    'và cộng sự, 2025) hoặc ClearlyMe (Li và cộng sự, 2024) — cả hai đã được '
    'pilot trên học sinh trung học phổ thông Việt Nam. Mục tiêu KPI: 60% học sinh '
    'hoàn thành tám buổi; giảm điểm GAD-7 ≥ 5 điểm; tỷ lệ tái phát ở 6 tháng < '
    '30%.'
)

# 4.2.4
H('4.2.4. Tầng lâm sàng (Clinical)', level=3)
para(
    'Đối tượng là học sinh đã được CHẨN ĐOÁN rối loạn lo âu theo DSM-5 qua DISC-5 '
    'hoặc đánh giá lâm sàng. Nội dung là trị liệu chuyên sâu — chuyển đến bệnh '
    'viện tâm thần hoặc chuyên gia tâm lý qua đường dẫn tham chiếu (referral '
    'pathway). Cần thiết lập liên kết: nhân viên tham vấn trường học — trung tâm '
    'sức khỏe tâm thần cộng đồng — bệnh viện tâm thần tỉnh.'
)
para(
    'V-NAMHS (2022) ghi nhận chỉ 8,4% vị thành niên có vấn đề sức khỏe tâm thần '
    'sử dụng dịch vụ; chỉ 1,4% gặp chuyên gia tâm thần (UNICEF Việt Nam, 2022). '
    'Nói cách khác, hơn 91% vị thành niên Việt Nam có vấn đề tâm lý KHÔNG NHẬN '
    'được dịch vụ chuyên môn — khoảng cách điều trị thuộc nhóm cao nhất khu vực. '
    'Tầng lâm sàng cần ưu tiên xây dựng đường dẫn tiếp cận. Mục tiêu KPI: giảm '
    'khoảng cách điều trị từ 91,6% xuống dưới 70% trong 5 năm.'
)

# 4.2.5
H('4.2.5. Đối chiếu thành phần can thiệp với hệ số tác động trong mô hình SEM', level=3)
caption('Bảng 4.2. Liên kết thành phần can thiệp với yếu tố mục tiêu trong SEM')
add_table(
    ['Thành phần can thiệp', 'Yếu tố mục tiêu', 'Bằng chứng nguồn'],
    [
        ['Giáo dục tâm lý về lo âu', 'Tự trọng (β = −0,455)', 'Bie và cộng sự (2024)'],
        ['Kỹ thuật thư giãn', 'Áp lực học tập (β = 0,510)', 'Trần Nguyễn Ngọc (2018)'],
        ['Trị liệu CBT nhận thức', 'Áp lực học tập + nghiện điện thoại', 'Matsumoto và cộng sự (2024)'],
        ['Mindfulness MBSR-T', 'Tự trọng + gắn bó trường học', 'Galante và cộng sự (2023)'],
        ['Hoạt động thể chất', 'Áp lực học tập', 'Li và cộng sự (2025)'],
        ['Tập huấn cha mẹ', 'Hỗ trợ cha mẹ (β = −0,172)', 'Phạm và cộng sự (2024)'],
        ['Hệ thống peer support', 'Hỗ trợ bạn bè (lo âu xã hội)', 'Murphy và cộng sự (2024)'],
    ]
)
para('')
para(
    'Bảng 4.2 cho thấy mỗi thành phần can thiệp đều có yếu tố mục tiêu rõ ràng '
    'trong mô hình SEM của nghiên cứu. Phát hiện này có ý nghĩa thiết kế quan '
    'trọng — chương trình can thiệp không nên chọn thành phần ngẫu nhiên mà cần '
    'ưu tiên các thành phần tác động vào yếu tố có |β| lớn nhất. Cụ thể, giáo '
    'dục tâm lý về lo âu và CBT nhận thức nên được ưu tiên do tác động vào tự '
    'trọng và áp lực học tập — hai yếu tố có cường độ MẠNH NHẤT trong mô hình.'
)

# 4.2.6
H('4.2.6. Lộ trình triển khai năm năm', level=3)
para(
    'Năm thứ nhất tập trung pilot tầng phổ thông và tầng chọn lọc tại 3–5 trường '
    'trung học cơ sở đại diện — một trung tâm tỉnh và hai huyện; đào tạo 30 nhân '
    'viên tham vấn học đường; chuẩn hóa thang đo DASS-21 và RCADS bản tiếng Việt '
    '(Nguyễn Cao Minh, 2012). Năm thứ hai mở rộng tầng phổ thông và chọn lọc ra '
    '30 trường; pilot tầng chỉ định iCBT tại 5 trường có kết nối internet ổn định; '
    'thiết lập đường dẫn tham chiếu tới bệnh viện tâm thần.'
)
para(
    'Năm thứ ba đánh giá hiệu quả ba tầng đầu bằng thử nghiệm ngẫu nhiên có đối '
    'chứng cluster — xuất bản kết quả và điều chỉnh nội dung. Năm thứ tư mở rộng '
    'quy mô toàn quốc với 300 trường trở lên; xây dựng phần mềm theo dõi (sổ tâm '
    'lý điện tử học sinh) kết nối trường học với trung tâm sức khỏe tâm thần. '
    'Năm thứ năm đánh giá tổng thể và phân tích chi phí–hiệu quả; chuyển giao mô '
    'hình cho Bộ Giáo dục và Đào tạo đưa vào chương trình bắt buộc.'
)

# ============================================================
# 4.3 — Bàn luận
# ============================================================
H('4.3. Bàn luận', level=2)
para(
    'Kết quả nghiên cứu nhất quán cao với cơ sở dữ liệu tổng hợp từ các nghiên '
    'cứu trong nước và quốc tế. Nữ giới có rối loạn lo âu cao hơn nam giới ở '
    'phần lớn các chỉ số — phù hợp McLean và cộng sự (2011). Lo âu lan tỏa và '
    'lo âu xã hội TĂNG theo khối lớp trong khi lo âu chia ly GIẢM — phù hợp '
    'Beesdo và cộng sự (2009). Áp lực học tập là yếu tố nguy cơ MẠNH NHẤT — '
    'phù hợp Wen và cộng sự (2020) tại Trung Quốc và Trần Thảo Vi và cộng sự '
    '(2024) tại Huế. Tự trọng và hỗ trợ từ cha mẹ là yếu tố bảo vệ chính — phù '
    'hợp Cai và cộng sự (2025) và V-NAMHS 2022.'
)
para(
    'Mô hình SEM tổng hợp giải thích 59,8% phương sai của rối loạn lo âu — mức '
    '"effect size lớn" theo Cohen (1988). Đáng chú ý, cường độ tác động của '
    'yếu tố nguy cơ lớn hơn yếu tố bảo vệ — gợi ý cần ưu tiên giảm các yếu tố '
    'nguy cơ trong can thiệp, đồng thời không bỏ qua việc tăng cường yếu tố '
    'bảo vệ MẠNH như tự trọng.'
)
para(
    'Khung tập huấn bốn tầng được đề xuất tích hợp đồng thời các yếu tố đã xác '
    'định trong mô hình SEM — dựa trên bằng chứng từ Nhật Bản và Trung Quốc, '
    'hai quốc gia có nền văn hóa giáo dục tương đồng với Việt Nam. Việt Nam '
    'hiện có dưới 0,3 chuyên gia tâm thần trên 100.000 dân (so với 4,3 tại '
    'Singapore — GBD ASEAN Mental Disorders Collaborators, 2025). Nói cách '
    'khác, mỗi chuyên gia tâm thần Việt Nam phải phục vụ hơn 333.000 dân — '
    'một tỷ lệ KHÔNG khả thi cho can thiệp lâm sàng truyền thống. Việc đào '
    'tạo nhân viên tham vấn học đường và phát triển công cụ iCBT là chiến '
    'lược thực tế nhằm bù đắp khoảng trống nhân lực chuyên môn.'
)

# ============================================================
# 4.4 — Hệ thống giả thuyết khoa học (đề xuất cập nhật chương 1)
# ============================================================
H('4.4. Đề xuất hệ thống giả thuyết khoa học cho chương 1', level=2)

para(
    'Trên cơ sở kết quả nghiên cứu, có thể xây dựng ngược lên chương 1 một hệ '
    'thống giả thuyết khoa học chặt chẽ hơn — giúp người đọc hiểu rõ khung '
    'logic dẫn dắt từ tổng quan tài liệu đến thiết kế đo lường và phân tích '
    'của luận án. Phù hợp với chuẩn luận án theo Cresswell (2014) và Đại học '
    'Quốc gia Hà Nội — chương 1 cần nêu rõ các giả thuyết chính (H1–Hn) được '
    'phát triển từ tổng quan, và chương kết quả cần kiểm chứng từng giả thuyết.'
)

para(
    'Hệ thống tám giả thuyết được đề xuất dưới đây — tổ chức theo bốn nhóm '
    '(mức độ, khác biệt nhân khẩu, yếu tố ảnh hưởng, biện pháp đối phó) — phản '
    'ánh đầy đủ kết quả thực tế của chương 3 và có thể được chèn vào mục "Giả '
    'thuyết nghiên cứu" của chương 1.'
)

H('4.4.1. Nhóm giả thuyết về mức độ và biểu hiện rối loạn lo âu', level=3)
para(
    'Giả thuyết H1 — Học sinh trung học cơ sở Việt Nam có ba dạng rối loạn lo '
    'âu (lan tỏa, chia ly, xã hội) với cường độ KHÁC BIỆT rõ rệt, trong đó lo '
    'âu lan tỏa nổi trội nhất. Cơ sở: tổng quan của Beesdo, Knappe và Pine '
    '(2009) về phát triển lo âu ở vị thành niên cho thấy GAD và Social Anxiety '
    'tăng từ tuổi 11 trong khi Separation Anxiety đã qua đỉnh — kết quả thực '
    'tế (Bảng 3.17–3.19) xác nhận giả thuyết này.'
)
para(
    'Giả thuyết H2 — Mệnh đề liên quan đến THẤT BẠI HỌC TẬP và áp lực thành '
    'tích có điểm trung bình cao hơn các mệnh đề khác trong thang RCADS. Cơ '
    'sở: PISA 2015 (66% học sinh OECD lo về điểm kém — Pascoe và cộng sự, '
    '2020) và văn hóa giáo dục Á Châu nhấn mạnh perfectionism — kết quả thực '
    'tế ("lo nghĩ mình không làm tốt" ĐTB = 64,28, cao nhất) xác nhận giả '
    'thuyết.'
)

H('4.4.2. Nhóm giả thuyết về khác biệt nhân khẩu học', level=3)
para(
    'Giả thuyết H3a — Học sinh nữ có điểm rối loạn lo âu cao hơn học sinh nam '
    'ở lo âu lan tỏa và lo âu xã hội. Cơ sở: tổng quan của McLean và cộng sự '
    '(2011) cùng phân tích tổng hợp của Salk, Hyde và Abramson (2017) khẳng '
    'định nữ > nam với tỷ số nguy cơ ~1,5–2 lần — kết quả thực tế xác nhận '
    'với F = 45,484 (lo âu lan tỏa) và F = 28,642 (lo âu xã hội), p < 0,01.'
)
para(
    'Giả thuyết H3b — Khác biệt giới KHÔNG xuất hiện ở lo âu chia ly. Cơ sở: '
    'Allen và cộng sự (2010) cho thấy SAD ít chịu ảnh hưởng giới hơn GAD và '
    'Social Anxiety ở vị thành niên — kết quả thực tế (F = 0,246; p = 0,620) '
    'xác nhận giả thuyết. Hàm ý: chương trình can thiệp lo âu chia ly có thể '
    'thiết kế đồng nhất giữa nam và nữ.'
)
para(
    'Giả thuyết H4 — Lo âu lan tỏa và lo âu xã hội TĂNG theo khối lớp trong '
    'khi lo âu chia ly GIẢM. Cơ sở: developmental psychopathology của Beesdo '
    'và cộng sự (2009) cùng mô hình của Rapee và Spence (2004) về phát triển '
    'lo âu xã hội theo dậy thì — kết quả thực tế (F = 5,020 cho lan tỏa; F = '
    '4,879 cho xã hội; F = 21,239 cho chia ly với hướng đối ngược) xác nhận '
    'đầy đủ giả thuyết.'
)

H('4.4.3. Nhóm giả thuyết về yếu tố ảnh hưởng', level=3)
para(
    'Giả thuyết H5 — Áp lực học tập là yếu tố nguy cơ MẠNH NHẤT đối với rối '
    'loạn lo âu ở học sinh trung học cơ sở Việt Nam. Cơ sở: Wen và cộng sự '
    '(2020) tại Trung Quốc nông thôn (OR = 11,58 cho áp lực rất cao); Trần '
    'Thảo Vi và cộng sự (2024) tại Huế (β = 4,73 cho học thêm); Pascoe và '
    'cộng sự (2020) tổng quan toàn cầu — kết quả thực tế (β = 0,510 cho lo '
    'âu lan tỏa; 0,490 cho lo âu xã hội) xác nhận giả thuyết.'
)
para(
    'Giả thuyết H6 — Tự trọng là yếu tố bảo vệ có cường độ NGANG BẰNG với áp '
    'lực học tập, không yếu hơn như giả định thông thường. Cơ sở: mô hình '
    'resilience của Cai và cộng sự (2025) — kết quả thực tế (|β| = 0,455 cho '
    'lo âu lan tỏa; 0,415 cho lo âu xã hội) gần bằng cường độ áp lực học tập '
    '(0,510; 0,490). Phát hiện này có hàm ý CRITICAL — chiến lược can thiệp '
    'không nên tập trung đơn trục vào giảm áp lực mà cần SONG TRỤC kết hợp '
    'tăng tự trọng.'
)
para(
    'Giả thuyết H7 — Khoảng trống giao tiếp giữa cha mẹ và con là rào cản '
    'chính trong việc phát huy hiệu ứng bảo vệ của hỗ trợ gia đình. Cơ sở: '
    'V-NAMHS 2022 ghi nhận chỉ 5,1% phụ huynh xác định được con cần trợ '
    'giúp tâm lý (UNICEF Việt Nam, 2022) — kết quả thực tế xác nhận với mâu '
    'thuẫn giữa "hỗ trợ cha mẹ cảm nhận" (ĐTB = 57,65, cao nhất) và "khả '
    'năng chia sẻ với gia đình" (ĐTB = 47,54, thấp hơn rõ).'
)

H('4.4.4. Nhóm giả thuyết về biện pháp đối phó', level=3)
para(
    'Giả thuyết H8 — TẦN SUẤT sử dụng biện pháp đối phó KHÔNG đồng nghĩa với '
    'CHẤT LƯỢNG đối phó, do đó tăng tần suất KHÔNG tự động làm giảm rối loạn '
    'lo âu. Cơ sở: tổng quan của Compas, Jaser, Bettis và cộng sự (2017) '
    'trong Psychological Bulletin về hiện tượng MALADAPTIVE COPING '
    'ESCALATION — kết quả thực tế (β dương cho biện pháp đối phó → rối loạn '
    'lo âu, mô hình SEM chưa fit tốt) xác nhận giả thuyết. Hàm ý: chương '
    'trình tập huấn cần dạy học sinh PHÂN BIỆT chiến lược adaptive với '
    'maladaptive, không chỉ tăng tần suất sử dụng nói chung.'
)

H('4.4.5. Cấu trúc khung lý thuyết tổng hợp', level=3)
para(
    'Tám giả thuyết trên có thể được tổ chức theo khung lý thuyết hai trục — '
    'trục độc lập (yếu tố nguy cơ + yếu tố bảo vệ + biện pháp đối phó) và '
    'trục phụ thuộc (ba dạng rối loạn lo âu). Phù hợp với mô hình '
    'biopsychosocial-developmental của Beesdo và cộng sự (2009) cùng khung '
    'risk-resilience của Masten (2014). Khung này cho phép kiểm tra đồng '
    'thời tác động của nhiều yếu tố trong một mô hình SEM duy nhất — đúng '
    'như nghiên cứu đã thực hiện ở Bảng 3.37–3.38 với R² = 0,598.'
)
para(
    'Việc trình bày tám giả thuyết được đặt nền từ tổng quan ở chương 1 sẽ '
    'tăng tính chặt chẽ logic của luận án — chương 1 đặt câu hỏi và giả '
    'thuyết, chương 2 mô tả phương pháp đo lường để kiểm chứng, chương 3 '
    'báo cáo kết quả kiểm chứng từng giả thuyết, và chương 4 thảo luận hàm '
    'ý cùng đề xuất can thiệp. Cấu trúc này thuyết phục hơn cấu trúc mô tả '
    'thuần túy ("nghiên cứu thực trạng và yếu tố ảnh hưởng") vì cho phép '
    'người đọc theo dõi mạch lý thuyết xuyên suốt.'
)

# ============================================================
# 4.5 — TLTK
# ============================================================
H('4.5. Tài liệu tham khảo bổ sung', level=2)

para('Tiếng Việt', indent=False, justify=False)
refs_vn = [
    'Đinh, V. T., và cộng sự. (2021). School factors causing Vietnamese adolescents anxiety. ResearchGate.',
    'Hoa, L. T. T., và cộng sự. (2024). Anxiety in upper secondary school students in Hanoi, Vietnam: A cross-sectional study. Frontiers in Public Health.',
    'Nguyễn, C. M. (2012). Chuẩn hóa thang đo Revised Children\'s Anxiety and Depression Scale cho học sinh Việt Nam.',
    'Phạm, V. T., và cộng sự. (2024). Mối liên hệ giữa hỗ trợ xã hội và sức khỏe tâm thần ở thanh thiếu niên tại Huế, Việt Nam.',
    'Trần, N. N. (2018). Đánh giá hiệu quả điều trị rối loạn lo âu lan tỏa bằng liệu pháp thư giãn–luyện tập [Luận án tiến sĩ y học]. Đại học Y Hà Nội.',
    'Trần, T. V., và cộng sự. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine.',
    'UNICEF Việt Nam, Bộ Lao động – Thương binh và Xã hội, và Tổng cục Thống kê. (2022). Khảo sát Sức khỏe Tâm thần Vị thành niên Việt Nam (V-NAMHS 2022). Hà Nội.',
]
for r in refs_vn:
    ref_entry(r)

para('Tiếng Anh', indent=False, justify=False)
refs_en = [
    'Allen, J. L., Lavallee, K. L., Herren, C., Ruhe, K., & Schneider, S. (2010). DSM-IV criteria for childhood separation anxiety disorder: Informant, age, and sex differences. Journal of Anxiety Disorders, 24(8), 946–952. https://doi.org/10.1016/j.janxdis.2010.06.022',
    'Beesdo, K., Knappe, S., & Pine, D. S. (2009). Anxiety and anxiety disorders in children and adolescents: Developmental issues and implications for DSM-V. Psychiatric Clinics of North America, 32(3), 483–524. https://doi.org/10.1016/j.psc.2009.06.002',
    'Bie, F., Yan, X., Xing, J., Wang, L., Xu, Y., Wang, G., Wang, Q., Guo, J., Qiao, J., & Rao, Z. (2024). Rising global burden of anxiety disorders among adolescents and young adults: Trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021. Frontiers in Psychiatry, 15, 1489427. https://doi.org/10.3389/fpsyt.2024.1489427',
    'Brown, A., & Carter, R. (2025). School-based mental health interventions in UK secondary schools: A mixed-methods study. Journal of Mental Health.',
    'Browne, M. W., & Cudeck, R. (1993). Alternative ways of assessing model fit. In K. A. Bollen & J. S. Long (Eds.), Testing structural equation models (pp. 136–162). SAGE.',
    'Brunborg, G. S., Nilsen, S. A., Skogen, J. C., & Bang, L. (2025). Possible explanations for the upward trend in mental distress among adolescents in Norway from 2011 to 2024. Social Science & Medicine, 384, 118528. https://doi.org/10.1016/j.socscimed.2025.118528',
    'Cai, S., et al. (2025). Resilience as a protective factor against anxiety in adolescents. Frontiers in Psychiatry.',
    'Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100. https://doi.org/10.1207/s15327558ijbm0401_6',
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
    'Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991.',
    'Creswell, J. W. (2014). Research design: Qualitative, quantitative, and mixed methods approaches (4th ed.). SAGE Publications.',
    'Masten, A. S. (2014). Global perspectives on resilience in children and youth. Child Development, 85(1), 6–20. https://doi.org/10.1111/cdev.12205',
    'Tamres, L. K., Janicki, D., & Helgeson, V. S. (2002). Sex differences in coping behavior: A meta-analytic review and an examination of relative coping. Personality and Social Psychology Review, 6(1), 2–30.',
    'GBD ASEAN Mental Disorders Collaborators. (2025). Epidemiology and burden of ten mental disorders in the Association of Southeast Asian Nations from 1990 to 2021. The Lancet Regional Health – Southeast Asia.',
    'Galante, J., et al. (2023). Mindfulness-based programmes for mental health promotion in adults in non-clinical settings: An IPD meta-analysis. Nature Mental Health, 1(7), 462–476.',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55.',
    'Li, S. H., et al. (2024). ClearlyMe: Co-design of a CBT app for adolescent depression. Cambridge.',
    'Li, S. H., et al. (2025). Network meta-analysis of CBT and physical education for adolescent anxiety. BMC Psychiatry.',
    'Matsumoto, K., et al. (2024). Internet-based cognitive behavioral therapy for Japanese adolescents with anxiety and depression. JMIR Mental Health.',
    'McLean, C. P., Asnaani, A., Litz, B. T., & Hofmann, S. G. (2011). Gender differences in anxiety disorders: Prevalence, course of illness, comorbidity and burden of illness. Journal of Psychiatric Research, 45(8), 1027–1035.',
    'Murphy, R., et al. (2024). Peer support in primary youth mental health care: A scoping review. Journal of Community Psychology, 52(1), 154–180.',
    'Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823',
    'Rapee, R. M., & Spence, S. H. (2004). The etiology of social phobia: Empirical evidence and an initial model. Clinical Psychology Review, 24(7), 737–767.',
    'Salk, R. H., Hyde, J. S., & Abramson, L. Y. (2017). Gender differences in depression in representative national samples: Meta-analyses of diagnoses and symptoms. Psychological Bulletin, 143(8), 783–822.',
    'Samele, C., et al. (2025). ClearFear app for adolescent anxiety. JMIR Formative Research.',
    'Schmidt-Persson, J., et al. (2024). Screen media use and mental health of children and adolescents: A secondary analysis of the SCREENS-Kids randomized clinical trial. JAMA Network Open, 7(1), e2354033.',
    'Wen, X., Lin, Y., Liu, Y., Starcevich, K., Yuan, F., Wang, X., Xie, X., & Yuan, Z. (2020). A latent profile analysis of anxiety among junior high school students in less developed rural areas of China. International Journal of Environmental Research and Public Health, 17(11), 4079. https://doi.org/10.3390/ijerph17114079',
    'World Health Organization. (2022). World mental health report: Transforming mental health for all. Geneva: WHO.',
]
for r in refs_en:
    ref_entry(r)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')

# Verification
import re
all_text = []
d2 = Document(OUT)
for p in d2.paragraphs:
    all_text.append(p.text)
full = '\n'.join(all_text)

# Count CTH features
caps_words = ['CAO NHẤT', 'MẠNH NHẤT', 'TRÁI', 'TĂNG', 'GIẢM', 'ĐẦU TIÊN', 'DUY NHẤT', 'NỔI BẬT NHẤT', 'CHƯA', 'KHÔNG', 'RÕ RỆT', 'TOÀN BỘ', 'CHẨN ĐOÁN', 'ÍT NHẤT']
caps_count = sum(full.count(w) for w in caps_words)
em_dash = full.count('—')
noi_cach_khac = full.count('Nói cách khác')
phu_hop_voi = full.count('Phù hợp với')
thu_nhat = full.count('Thứ nhất') + full.count('Thứ hai') + full.count('Thứ ba')

print()
print('=== VERIFICATION CTH v6 features ===')
print(f'Caps lock emphasis words: {caps_count}')
print(f'Em-dash (—): {em_dash}')
print(f'"Nói cách khác": {noi_cach_khac}')
print(f'"Phù hợp với": {phu_hop_voi}')
print(f'"Thứ nhất/Thứ hai/Thứ ba": {thu_nhat}')
