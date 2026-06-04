# -*- coding: utf-8 -*-
"""
Script v4: Docx with cropped & enhanced figures/tables from PDF,
inserted at correct positions in Vietnamese translation.
"""
import os, glob, sys
import fitz  # PyMuPDF
import cv2
import numpy as np
from PIL import Image, ImageEnhance, ImageFilter
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
TRANS_DIR = os.path.join(BASE, "Translations")
ORIG_DIR = os.path.join(BASE, "Originals")
OUT_DIR = os.path.join(BASE, "DocFiles")
CROP_DIR = os.path.join(BASE, "Cropped_Figures")
os.makedirs(OUT_DIR, exist_ok=True)
os.makedirs(CROP_DIR, exist_ok=True)


def crop_and_enhance(pdf_path, page_num, rect, output_name, zoom=3.0):
    """Crop a region from PDF page, enhance quality, save as PNG."""
    try:
        doc = fitz.open(pdf_path)
        page = doc[page_num - 1]  # 0-indexed
        # Crop region with high zoom
        mat = fitz.Matrix(zoom, zoom)
        clip = fitz.Rect(rect)
        pix = page.get_pixmap(matrix=mat, clip=clip)
        raw_path = os.path.join(CROP_DIR, f"{output_name}_raw.png")
        pix.save(raw_path)
        doc.close()

        # Enhance with Pillow
        img = Image.open(raw_path)
        # Sharpen
        img = img.filter(ImageFilter.SHARPEN)
        # Increase contrast
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.3)
        # Increase sharpness more
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.5)
        # Brightness slight boost
        enhancer = ImageEnhance.Brightness(img)
        img = enhancer.enhance(1.05)

        final_path = os.path.join(CROP_DIR, f"{output_name}.png")
        img.save(final_path, dpi=(300, 300))
        os.remove(raw_path)
        return final_path
    except Exception as e:
        print(f"    Error cropping {output_name}: {e}")
        return None


def extract_full_page_enhanced(pdf_path, page_num, output_name, zoom=3.0):
    """Extract full page at high resolution + enhance."""
    try:
        doc = fitz.open(pdf_path)
        page = doc[page_num - 1]
        mat = fitz.Matrix(zoom, zoom)
        pix = page.get_pixmap(matrix=mat)
        raw_path = os.path.join(CROP_DIR, f"{output_name}_raw.png")
        pix.save(raw_path)
        doc.close()

        img = Image.open(raw_path)
        img = img.filter(ImageFilter.SHARPEN)
        enhancer = ImageEnhance.Contrast(img)
        img = enhancer.enhance(1.2)
        enhancer = ImageEnhance.Sharpness(img)
        img = enhancer.enhance(1.4)

        final_path = os.path.join(CROP_DIR, f"{output_name}.png")
        img.save(final_path, dpi=(300, 300))
        os.remove(raw_path)
        return final_path
    except Exception as e:
        print(f"    Error: {e}")
        return None


# Define figures/tables to crop for each paper
# Format: {paper: [(label, pdf_page, rect_or_"full", insert_after_section), ...]}
# rect = (x0, y0, x1, y1) in PDF points (72 pts = 1 inch)
# "full" = extract full page

FIGURES = {
    "01_Jenkins_et_al_2023": [
        ("Table 1: Sociodemographic characteristics", 4, "full", "KET QUA"),
        ("Table 2: PHQ-9A and GAD-10 responses", 4, "full", "KET QUA"),
        ("Figure 1: PHQ-9A and GAD-10 distributions", 5, (30, 30, 570, 450), "KET QUA"),
        ("Figure 2: Gender differences box plots", 5, (30, 450, 570, 790), "KET QUA"),
    ],
    "02_Saikia_et_al_2023": [
        ("Table 1: Depression risk factors", 3, "full", "KET QUA"),
        ("Table 2: Anxiety risk factors", 4, "full", "KET QUA"),
        ("Table 3: Stress risk factors", 5, "full", "KET QUA"),
    ],
    "03_Mandaknalli_Malusare_2021": [
        ("Page 2 - Results", 2, "full", "KET QUA"),
        ("Page 3 - Tables", 3, "full", "KET QUA"),
    ],
    "05_Alharbi_et_al_2019": [
        ("Figure 1: Depression distribution PHQ-9", 2, (30, 520, 420, 730), "KET QUA"),
        ("Figure 2: Anxiety distribution GAD-7", 3, (30, 30, 420, 180), "KET QUA"),
        ("Table 1: Socio-demographics", 3, (30, 180, 420, 750), "KET QUA"),
        ("Table 2: Depression vs socio-demographics", 4, "full", "KET QUA"),
        ("Table 3: Anxiety vs socio-demographics", 5, "full", "KET QUA"),
    ],
    "06_Nakie_et_al_2022": [
        ("Table 1: Socio-demographics", 5, "full", "KET QUA"),
        ("Figure 1: DAS prevalence", 4, (280, 400, 570, 700), "KET QUA"),
        ("Table 4: Depression factors", 6, "full", "KET QUA"),
        ("Table 5: Anxiety factors", 7, "full", "KET QUA"),
        ("Table 6: Stress factors", 8, "full", "KET QUA"),
    ],
    "07_Chen_et_al_2023": [
        ("Table 1: Characteristics", 3, "full", "KET QUA"),
        ("Table 2: Prevalence", 4, "full", "KET QUA"),
        ("Table 3: Risk factors depression", 5, "full", "KET QUA"),
        ("Table 4: Risk factors anxiety", 6, "full", "KET QUA"),
    ],
    "08_Wen_et_al_2020": [
        ("Table 1: Demographics", 5, "full", "KET QUA"),
        ("Table 2: LPA fit indices", 6, "full", "KET QUA"),
        ("Figure 1: Profile plot", 7, (30, 30, 570, 400), "KET QUA"),
        ("Table 3: Logistic regression", 8, "full", "KET QUA"),
    ],
    "09_Qiu_et_al_2022": [
        ("Table 1: Characteristics", 4, (30, 30, 570, 480), "KET QUA"),
        ("Table 2: LPA fit indices", 4, (30, 480, 570, 600), "KET QUA"),
        ("Figure 1: Parenting profiles", 5, (30, 30, 570, 320), "KET QUA"),
        ("Table 3: Parenting & resilience", 5, (30, 320, 570, 750), "KET QUA"),
    ],
    "10_Xu_et_al_2021": [
        ("Table 1: Characteristics by anxiety", 2, (280, 30, 570, 750), "KET QUA"),
        ("Figure 1: Prevalence by location", 3, (30, 30, 570, 370), "KET QUA"),
        ("Figure 2: COVID cognition", 3, (30, 520, 570, 780), "KET QUA"),
        ("Table 2: Anxiety severity", 4, (30, 30, 570, 200), "KET QUA"),
        ("Figure 3: GAD symptoms", 4, (30, 380, 570, 750), "KET QUA"),
        ("Table 3: Logistic regression", 5, "full", "KET QUA"),
    ],
    "11_Bhardwaj_et_al_2020": [
        ("Page 6 - Results tables", 6, "full", "KET QUA"),
        ("Page 7 - Results tables", 7, "full", "KET QUA"),
        ("Page 8 - Results figures", 8, "full", "KET QUA"),
        ("Page 9 - Results figures", 9, "full", "KET QUA"),
    ],
}


def add_word_table(doc, table_lines, in_critical=False):
    data_lines = []
    for tl in table_lines:
        tl = tl.strip()
        if tl and not all(c in '|-: ' for c in tl):
            cells = [c.strip() for c in tl.strip('|').split('|')]
            data_lines.append(cells)
    if not data_lines:
        return
    num_cols = max(len(row) for row in data_lines)
    for row in data_lines:
        while len(row) < num_cols:
            row.append('')
    table = doc.add_table(rows=len(data_lines), cols=num_cols)
    table.style = 'Table Grid'
    for i, row_data in enumerate(data_lines):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            clean = cell_text.strip()
            if '**' in clean:
                parts = clean.split('**')
                for k, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        if k % 2 == 1: run.bold = True
                        if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
            else:
                run = p.add_run(clean)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
            if i == 0:
                for run in p.runs: run.bold = True
    doc.add_paragraph()


def insert_figure(doc, img_path, caption):
    """Insert enhanced figure with caption."""
    try:
        doc.add_picture(img_path, width=Inches(6.2))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()
    except Exception as e:
        doc.add_paragraph(f"[Error inserting figure: {e}]")


def build_docx(md_path, docx_path, figures_list=None):
    """Build docx from markdown with figures inserted after sections."""
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_critical = False
    table_buf = []
    current_section = ""
    fig_inserted = set()

    def flush_figures_for(section_keyword):
        if not figures_list:
            return
        for i, (label, _, _, insert_after) in enumerate(figures_list):
            if i not in fig_inserted and insert_after.upper() in section_keyword.upper():
                img_path = figures_list[i][1]  # will be replaced with actual path
                # We store paths in a separate dict
                pass

    for line in lines:
        line = line.rstrip('\n')

        # Table collection
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            table_buf.append(line)
            continue
        elif table_buf:
            add_word_table(doc, table_buf, in_critical)
            table_buf = []

        if not line.strip():
            continue

        if 'QUAN DIEM PHAN BIEN' in line.upper() or 'CRITICAL REVIEW' in line:
            # Insert remaining figures before critical review
            if figures_list:
                for i, (label, img_path, _, _) in enumerate(figures_list):
                    if i not in fig_inserted and img_path and os.path.exists(img_path):
                        insert_figure(doc, img_path, label)
                        fig_inserted.add(i)
            in_critical = True

        # Section change -> insert figures for previous section
        if line.startswith('## ') and figures_list:
            for i, (label, img_path, _, insert_kw) in enumerate(figures_list):
                if i not in fig_inserted and img_path and os.path.exists(img_path):
                    if current_section and insert_kw.upper() in current_section.upper():
                        insert_figure(doc, img_path, label)
                        fig_inserted.add(i)
            current_section = line

        # Render
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            if in_critical:
                for r in h.runs: r.font.color.rgb = RGBColor(0xCC, 0, 0)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('---'):
            pass
        elif line.lstrip().startswith('- '):
            text = line.lstrip(' -')
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                for k, part in enumerate(text.split('**')):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                        if k % 2 == 1: run.bold = True
                        if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
        else:
            p = doc.add_paragraph()
            if '**' in line:
                for k, part in enumerate(line.split('**')):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(12)
                        if k % 2 == 1: run.bold = True
                        if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
            else:
                run = p.add_run(line)
                run.font.size = Pt(12)
                if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)

    if table_buf:
        add_word_table(doc, table_buf, in_critical)

    doc.save(docx_path)


# ===== MAIN =====
print("=" * 60)
print("DOCX V4: Crop + Enhance figures from PDF")
print("=" * 60)

pdf_files = {
    "01_Jenkins_et_al_2023": "01_Jenkins_et_al_2023.pdf",
    "02_Saikia_et_al_2023": "02_Saikia_et_al_2023.pdf",
    "03_Mandaknalli_Malusare_2021": "03_Mandaknalli_Malusare_2021.pdf",
    "05_Alharbi_et_al_2019": "05_Alharbi_et_al_2019.pdf",
    "06_Nakie_et_al_2022": "06_Nakie_et_al_2022.pdf",
    "07_Chen_et_al_2023": "07_Chen_et_al_2023.pdf",
    "08_Wen_et_al_2020": "08_Wen_et_al_2020.pdf",
    "09_Qiu_et_al_2022": "09_Qiu_et_al_2022.pdf",
    "10_Xu_et_al_2021": "10_Xu_et_al_2021.pdf",
    "11_Bhardwaj_et_al_2020": "11_Bhardwaj_et_al_2020.pdf",
}

# Step 1: Crop and enhance all figures
print("\n--- Cropping & enhancing figures ---")
all_figures = {}  # {paper_key: [(label, img_path, rect, insert_kw), ...]}

for paper_key, fig_list in FIGURES.items():
    pdf_name = pdf_files.get(paper_key)
    if not pdf_name:
        continue
    pdf_path = os.path.join(ORIG_DIR, pdf_name)
    if not os.path.exists(pdf_path):
        continue

    paper_figs = []
    for i, (label, page_num, rect, insert_kw) in enumerate(fig_list):
        out_name = f"{paper_key}_fig{i+1}"
        if rect == "full":
            img_path = extract_full_page_enhanced(pdf_path, page_num, out_name, zoom=3.0)
        else:
            img_path = crop_and_enhance(pdf_path, page_num, rect, out_name, zoom=3.0)
        paper_figs.append((label, img_path, rect, insert_kw))
        if img_path:
            print(f"  {out_name}: OK")
    all_figures[paper_key] = paper_figs

# Step 2: Build all docx
print("\n--- Building DOCX files ---")
for md_file in sorted(glob.glob(os.path.join(TRANS_DIR, "*.md"))):
    basename = os.path.splitext(os.path.basename(md_file))[0]
    docx_path = os.path.join(OUT_DIR, f"{basename}.docx")
    figs = all_figures.get(basename)
    build_docx(md_file, docx_path, figs)
    print(f"  OK: {basename}.docx")

# Summary
print("\n--- Summary ---")
summary_md = os.path.join(BASE, "BANG_TOM_TAT_11_BAI_BAO.md")
summary_docx = os.path.join(OUT_DIR, "BANG_TOM_TAT_TONG_HOP.docx")
build_docx(summary_md, summary_docx)
print("  OK: BANG_TOM_TAT_TONG_HOP.docx")

print(f"\n{'='*60}")
print("HOAN TAT!")
print(f"{'='*60}")
