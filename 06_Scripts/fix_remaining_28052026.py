# -*- coding: utf-8 -*-
"""Fix cac van de con sot:
1. 3 orphan Heading 2 (text trong) - revert ve Normal
2. PHỤ LỤC 1/2/3 + chapter title cua phụ lục - apply Heading 1
3. Para 771 (continuation cua CHƯƠNG 3 title) - apply Heading 1

28/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

# 1. Revert orphan headings
orphans = [214, 295, 346]
for idx in orphans:
    if idx >= len(d.paragraphs): continue
    p = d.paragraphs[idx]
    if p.style.name in ('Heading 1', 'Heading 2', 'Heading 3', 'Heading 4', 'Heading 5') and not p.text.strip():
        try:
            p.style = d.styles['Normal']
            print(f"  Reverted para {idx}: empty Heading 2 -> Normal")
        except: pass

# 2. Apply Heading 1 to PHỤ LỤC 1/2/3
phu_luc_paras = [1395, 1491, 1536]
for idx in phu_luc_paras:
    if idx >= len(d.paragraphs): continue
    p = d.paragraphs[idx]
    text = p.text.strip()
    if 'PHỤ LỤC' in text.upper():
        try:
            p.style = d.styles['Heading 1']
            print(f"  Applied Heading 1: para {idx}: {text}")
        except: pass

# 3. Apply Heading 1 to chapter title continuation (para 771)
if 771 < len(d.paragraphs):
    p = d.paragraphs[771]
    text = p.text.strip()
    if text.upper() == 'Ở HỌC SINH TRUNG HỌC CƠ SỞ':
        try:
            p.style = d.styles['Heading 1']
            print(f"  Applied Heading 1: para 771: {text} (continuation cua CHƯƠNG 3 title)")
        except: pass

d.save(FILE)
print(f"\nSaved: {FILE}")
