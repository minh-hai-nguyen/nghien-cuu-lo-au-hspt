"""Build VN020 FIXED tóm tắt — expanded from verified PDF facts."""
import os
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

OUT = r"c:\Users\OS\OneDrive\read_books\Lo-au\Tom-tat-tung-bai\VN020_TranHoVinhLoc_2024_TPHCM_DASS-Y_FIXED_07062026.docx"

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.0)

# Default style
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
rpr = style.element.rPr
if rpr is None:
    from docx.oxml import OxmlElement
    rpr = OxmlElement('w:rPr')
    style.element.append(rpr)
rfonts = style.element.rPr.find(qn('w:rFonts'))
if rfonts is None:
    from docx.oxml import OxmlElement
    rfonts = OxmlElement('w:rFonts')
    style.element.rPr.append(rfonts)
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')
rfonts.set(qn('w:cs'), 'Times New Roman')
rfonts.set(qn('w:eastAsia'), 'Times New Roman')

def set_spacing(p):
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.space_after = Pt(6)

def add_heading(text, level=1, bold=True, size=14):
    p = doc.add_paragraph()
    set_spacing(p)
    run = p.add_run(text)
    run.bold = bold
    run.font.name = 'Times New Roman'
    run.font.size = Pt(size)
    return p

def add_para(text, bold=False, italic=False):
    p = doc.add_paragraph()
    set_spacing(p)
    p.paragraph_format.first_line_indent = Cm(0.75)
    run = p.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    run.bold = bold
    run.italic = italic
    return p

def add_label_para(label, body):
    p = doc.add_paragraph()
    set_spacing(p)
    p.paragraph_format.first_line_indent = Cm(0.75)
    r1 = p.add_run(label)
    r1.bold = True
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(12)
    r2 = p.add_run(body)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(12)
    return p

# === TITLE ===
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(title)
r = title.add_run("TÓM TẮT BÀI VN020")
r.bold = True
r.font.size = Pt(14)
r.font.name = 'Times New Roman'

# === THÔNG TIN CÔNG TRÌNH ===
add_heading("1. Thông tin công trình", size=13)
add_label_para(
    "Tên bài: ",
    "Trầm cảm, lo âu, căng thẳng và các yếu tố liên quan ở học sinh trung học phổ thông tại Thành phố Hồ Chí Minh "
    "(Stress, Anxiety, Depression and Associated Factors among High School Students in Ho Chi Minh City)."
)
add_label_para(
    "Tác giả: ",
    "Trần Hồ Vĩnh Lộc¹, Huỳnh Ngọc Vân Anh²* (tác giả liên hệ — hnvanhytcc@ump.edu.vn), Tô Gia Kiên³. "
    "Đơn vị: ¹Khoa Y tế Công cộng, ²Bộ môn Thống kê và Tin học y học, ³Bộ môn Quản lý y tế — Đại học Y Dược TPHCM."
)
add_label_para(
    "Nguồn đăng: ",
    "Tạp chí Y học Thành phố Hồ Chí Minh, Tập 27, Số 5 (2024), trang 100–110. "
    "ISSN: 1859-1779. DOI: 10.32895/hcjm.m.2024.05.12."
)
add_label_para(
    "Mốc thời gian: ",
    "Nhận bài 04/10/2024 — chấp nhận 11/11/2024 — đăng 13/11/2024. "
    "Thu thập dữ liệu thực địa từ tháng 02/2024 đến tháng 04/2024. "
    "Đã được Hội đồng Đạo đức ĐH Y Dược TPHCM thông qua (số 276/HĐĐĐ-ĐHYD, ngày 01/02/2024)."
)
add_label_para(
    "Địa bàn — đối tượng: ",
    "Học sinh khối 10, 11, 12 tại hai trường THPT công lập của TPHCM: Trường THPT Bình Hưng Hòa (quận Bình Tân, 555 HS) "
    "và Trường THPT Đa Phước (huyện Bình Chánh, 431 HS). Tổng phát phiếu: 986 HS — 10 HS bị loại do trả lời không đầy đủ — "
    "N phân tích cuối cùng = 976 HS (99,0%)."
)

# === PHƯƠNG PHÁP ===
add_heading("2. Phương pháp nghiên cứu", size=13)
add_para(
    "Thiết kế cắt ngang mô tả, phân tích các yếu tố liên quan. Chọn mẫu nhiều bậc: ngẫu nhiên đơn 1 quận (Bình Tân) "
    "và 1 huyện (Bình Chánh) ở bậc đầu; ngẫu nhiên chọn 1 trường THPT/quận-huyện; ở bậc lớp dùng chọn thuận tiện 4 lớp/khối "
    "(khối 10, 11, 12), tất cả HS trong lớp được mời tham gia. Cỡ mẫu tối thiểu ước lượng 812 HS (k=2; từ chối 5%) "
    "dựa theo tỷ lệ TC/LÂ/CT 55,8%–67,9%–50,2% của Võ Minh Đức (2023). Nghiên cứu pilot trên 34 HS lớp 10C17 — "
    "Cronbach's α: DASS-Y = 0,94; ESSA = 0,74."
)
add_para(
    "Công cụ đo: (i) DASS-Y — Depression Anxiety Stress Scales for Youth — phiên bản chuyên biệt cho thanh thiếu niên (KHÔNG phải "
    "DASS-21 người lớn), gồm 21 mục thang Likert 4 mức (0 — Không đúng → 3 — Rất đúng). Ngưỡng cắt vị thành niên (Psychology "
    "Foundation of Australia, 2023): stress ≥ 12, lo âu ≥ 6, trầm cảm ≥ 7. (ii) ESSA — Educational Stress Scale for Adolescents "
    "(Sun et al., 2011) — 16 mục Likert 5 mức (1–5), tổng 16–80 điểm; phân nhóm: nhẹ < 51, vừa 51–58, nặng ≥ 59. "
    "Số liệu nhập bằng EpiData Entry Client 4.6.0.2 và phân tích bằng Stata 16.0. "
    "Mô hình hồi quy Poisson đa biến với tùy chọn Robust được dùng để ước lượng PR (Prevalence Ratio) thô và hiệu chỉnh, có hiệu chỉnh "
    "đồng thời các biến: khối lớp, giới tính, áp lực học tập, tình trạng hôn nhân cha mẹ, tình trạng sống chung của HS, kinh tế gia đình, "
    "số con, thứ tự con, nghề nghiệp cha-mẹ, học vấn cha-mẹ. Ngưỡng có ý nghĩa thống kê p < 0,05."
)

# === KẾT QUẢ ===
add_heading("3. Kết quả chính", size=13)
add_para(
    "Đặc điểm mẫu: tuổi trung bình 17,0 ± 0,8; nữ chiếm 53,4% (521/976); 82,5% HS sống cùng cả cha lẫn mẹ; 14,1% có cha mẹ "
    "đã ly thân/ly hôn; 3,4% có cha hoặc mẹ đã mất; 8,4% nhận thấy kinh tế gia đình khó khăn; mẹ có trình độ THCS chiếm tỷ lệ "
    "cao nhất (29,9%)."
)
add_para(
    "Tỷ lệ ba vấn đề SKTT (theo DASS-Y, mức nhẹ trở lên): trầm cảm 31,7% (n=309; KTC 95%: 28,7–34,7%); lo âu 25,1% (n=245; "
    "KTC 95%: 22,4–27,9%); căng thẳng 23,8% (n=232; KTC 95%: 21,1–26,6%). 42,4% (n=414) có ít nhất một vấn đề; 13,2% (n=129) "
    "đồng thời cả ba vấn đề. Phân bố theo mức độ trầm cảm: nhẹ 9,3%, trung bình 12,7%, nặng 5,4%, rất nặng 4,3% — cảnh báo "
    "một nhóm đáng kể (~9,7%) ở mức nặng/rất nặng cần can thiệp lâm sàng."
)
add_para(
    "Mô hình hồi quy Poisson đa biến (PR hiệu chỉnh, KTC 95%). Trầm cảm liên quan đến: nữ giới (PR=1,24; 1,03–1,50); áp lực "
    "học tập vừa (PR=1,82; 1,60–2,06) và nặng (PR=3,31; 2,57–4,27); cha mẹ ly thân/ly hôn (PR=1,44; 1,16–1,77); kinh tế đủ "
    "sống bảo vệ (PR=0,70; 0,54–0,88) và khá giả bảo vệ (PR=0,69; 0,50–0,95). Lo âu liên quan đến: nữ giới (PR=1,30; 1,03–1,63); "
    "áp lực học tập vừa (PR=1,68; 1,45–1,94) và nặng (PR=2,82; 2,11–3,77); mẹ học THPT giảm nguy cơ (PR=0,57; 0,36–0,88). "
    "Căng thẳng liên quan đến: nữ giới (PR=1,47; 1,15–1,87); áp lực học tập vừa (PR=1,92; 1,64–2,25) và nặng (PR=3,70; 2,71–5,06); "
    "có tôn giáo (PR=1,23; 1,01–1,52); cha mẹ ly thân/ly hôn (PR=1,42; 1,09–1,86); mẹ học vấn từ tiểu học → đại học đều có PR "
    "tăng (3,18–3,80) so với mẹ chỉ biết đọc, biết viết."
)

# === ĐỐI CHIẾU ===
add_heading("4. Đối chiếu liên bài & nhận định", size=13)
add_para(
    "Lo âu 25,1% (DASS-Y) thấp hơn rõ rệt so với các NC VN dùng DASS-21 trên cùng đối tượng THPT: Hoa 2024 (40,6% — GAD-7), "
    "Bảo Quyên 2025 (86,2% — DASS-21), Thảo Vi 2025 (65,8% — DASS-21), An Giang 2025 (61,2% — DASS-21), Vũ Thị Ly Ngọc 2018 "
    "(67,3% — DASS-21). Các tác giả lý giải khoảng cách này bằng đặc điểm tâm trắc của DASS-Y: ngưỡng cắt được điều chỉnh phù hợp "
    "đặc điểm phát triển VTN, mục từ đơn giản hơn, ít quy gán tình trạng bệnh lý quá mức như khi áp DASS-21 người lớn cho VTN "
    "(Cao C, 2023). Đây là phát hiện phương pháp luận đáng chú ý, gợi ý cộng đồng nghiên cứu SKTT VTN VN cần thận trọng khi so "
    "sánh tỷ lệ giữa các nghiên cứu dùng công cụ khác nhau."
)
add_para(
    "Áp lực học tập (ESSA) nổi lên là yếu tố nguy cơ mạnh nhất — PR nặng vs nhẹ: 3,31 (TC) / 2,82 (LÂ) / 3,70 (CT). Phù hợp Wen "
    "(2020 — OR áp lực = 11,6) và Phạm Minh Quang (2023). Cấu trúc gia đình (cha mẹ ly thân/ly hôn) làm tăng TC và CT, đồng nhất "
    "với Marouane M (2023 — OR = 6,00). Kinh tế gia đình đủ sống → khá giả có vai trò bảo vệ chống TC. Nữ > nam ở cả 3 chỉ số "
    "— phù hợp xu hướng toàn cầu (Kuehner 2017) liên quan dao động estrogen-progesterone và khác biệt cấu trúc não bộ "
    "(Luders & Toga 2010)."
)

# === KẾT LUẬN ===
add_heading("5. Kết luận của tác giả", size=13)
add_para(
    "Tỷ lệ TC/LÂ/CT ở học sinh THPT TPHCM (theo DASS-Y) lần lượt là 31,7% / 25,1% / 23,8% — cao một cách đáng quan ngại, với "
    "13,2% chịu đồng thời cả ba vấn đề. Các yếu tố liên quan đáng kể gồm giới tính, áp lực học tập, tình trạng hôn nhân của cha mẹ, "
    "kinh tế gia đình và trình độ học vấn của mẹ. Nhà trường cần điều chỉnh khối lượng bài tập, giảm áp lực thành tích, tổ chức tư vấn "
    "tâm lý và giáo dục SKTT cho HS và gia đình, đặc biệt là chương trình dành cho HS nữ."
)

# === PHẢN BIỆN ===
add_heading("6. Phản biện", size=13)
add_para(
    "Điểm mạnh: (i) Đơn vị thực hiện (ĐH Y Dược TPHCM) có uy tín; thông qua Hội đồng Đạo đức số hiệu cụ thể. (ii) Sử dụng DASS-Y "
    "— bản dành riêng cho VTN với ngưỡng cắt phù hợp — là đóng góp phương pháp luận hiếm gặp trong nhóm NC SKTT VTN VN. "
    "(iii) Cỡ mẫu n = 976 thuộc loại lớn so với mặt bằng các NC THPT tại TPHCM; tỷ lệ trả lời cao (99,0%). (iv) Cronbach's α của "
    "công cụ trong pilot ở mức rất tốt (DASS-Y = 0,94; ESSA = 0,74). (v) Hồi quy Poisson với Robust SE phù hợp với biến kết cục "
    "phổ biến (PR thay vì OR — tránh phóng đại liên kết)."
)
add_para(
    "Hạn chế: (i) Chỉ 2 trường công lập ở Bình Tân/Bình Chánh — chưa đại diện cho HS THPT TPHCM (thiếu trường tư, trường chuyên, "
    "trung tâm thành phố như Q1, Q3, Q5…). (ii) Bậc lớp dùng chọn thuận tiện — có thể gây sai lệch. (iii) Thiết kế cắt ngang — "
    "không suy diễn nhân quả. (iv) DASS-Y mới, dữ liệu chuẩn hóa cho VN còn hạn chế — cần thêm NC kiểm định tâm trắc trên mẫu VN. "
    "(v) Thiếu các biến quan trọng đã được tài liệu chứng minh: thời lượng/loại hình sử dụng MXH, chất lượng giấc ngủ, hoạt động thể "
    "chất, lịch sử tâm thần gia đình, bắt nạt học đường. (vi) Tạp chí Y học TPHCM chưa được PubMed indexed — hạn chế khả năng đối "
    "chiếu quốc tế. (vii) Thu thập sau Tết (02–04/2024) — tác giả tự nhận giúp giảm bias áp lực thi cử, nhưng cũng có thể bỏ sót giai đoạn căng nhất."
)

# === HƯỚNG TIẾP THEO ===
add_heading("7. Hướng nghiên cứu tiếp theo", size=13)
add_para(
    "(i) So sánh trực tiếp DASS-Y vs DASS-21 vs GAD-7/PHQ-A trên cùng mẫu VTN VN — xác định công cụ tối ưu, đề xuất ngưỡng cắt "
    "phù hợp VTN VN. (ii) Mở rộng sang nhiều trường (công lập-tư thục-chuyên), nhiều khu vực địa lý (nội thành-ngoại thành-nông "
    "thôn). (iii) Bổ sung biến MXH, giấc ngủ, hoạt động thể chất, bắt nạt — kiểm soát đầy đủ confounder. (iv) Thiết kế dọc/can thiệp: "
    "thử nghiệm chương trình giảm áp lực học tập, hỗ trợ HS có cha mẹ ly hôn, can thiệp cho HS nữ. (v) Đánh giá kết quả LA chuẩn hóa "
    "VN của DASS-Y với mẫu lớn đa trung tâm."
)
add_para("Đánh giá tổng hợp: ⭐⭐⭐ Trung bình–Khá. TC VN có DOI, công cụ DASS-Y mới, n = 976 — nhưng chỉ 2 trường, thiếu biến quan trọng.")

# === THAM KHẢO ===
add_heading("8. Tham khảo", size=13)
add_para(
    "PDF gốc: 02_Papers-goc/Viet-Nam/VN020_TranHoVinhLoc_2024_TPHCM_DASS-Y.pdf (10 trang, bản đăng chính thức Tập 27, Số 5, 2024)."
)
add_para(
    "Nguồn truy cập: Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh, Tô Gia Kiên. Trầm cảm, lo âu, căng thẳng và các yếu tố liên quan ở học "
    "sinh trung học phổ thông tại Thành phố Hồ Chí Minh. Tạp chí Y học TPHCM. 2024;27(5):100–110. DOI: 10.32895/hcjm.m.2024.05.12. "
    "URL: https://doi.org/10.32895/hcjm.m.2024.05.12 — https://www.tapchiyhoctphcm.vn."
)
add_para(
    "Tài liệu phương pháp luận trích dẫn trong bài gốc: Sun J, Dunne MP, Hou X, Xu A. Educational Stress Scale for Adolescents. "
    "Journal of Psychoeducational Assessment. 2011;29(6):534–46. — Psychology Foundation of Australia. DASS-Y severity cutoffs. 2023. "
    "https://dass.psy.unsw.edu.au/DASS-Y%20cutoffs.htm. — Cao C, Liao X, Gamble JH, et al. Evaluating the psychometric properties of the "
    "Chinese DASS-Y and DASS-21. Child Adolesc Psychiatry Ment Health. 2023;17(1):106."
)
add_para(
    "Truy vết nội bộ: canonical_index.json — id: VN020, descriptor: TranHoVinhLoc_2024_TPHCM_DASS-Y; thư mục pdf: Viet-Nam; "
    "đối chiếu các tóm tắt liên quan trong cùng cụm THPT TPHCM: Hoa 2024, Bảo Quyên 2025, Thảo Vi 2025, Vũ Thị Ly Ngọc 2018, "
    "Võ Minh Đức 2023."
)

# === FOOTER ===
footer = doc.add_paragraph()
footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
set_spacing(footer)
fr = footer.add_run("Đã sửa lại từ PDF gốc (07/06/2026)")
fr.italic = True
fr.font.name = 'Times New Roman'
fr.font.size = Pt(11)

# Strip metadata
cp = doc.core_properties
cp.author = ""
cp.last_modified_by = ""
cp.comments = ""
cp.keywords = ""
cp.subject = ""
cp.title = ""
cp.category = ""

doc.save(OUT)

# Word count
all_text = "\n".join(p.text for p in doc.paragraphs)
wc = len(all_text.split())
print(f"SAVED: {OUT}")
print(f"WORD COUNT: {wc}")
