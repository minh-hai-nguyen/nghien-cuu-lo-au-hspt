# -*- coding: utf-8 -*-
"""Audit kho ban dich + tom tat - tim cac loi tuong tu LA da phat hien.
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

FILES_TO_CHECK = [
    ('03_Ban-dich/VN016_BaoQuyen_2025_YHCD.docx', 'BAN DICH VN016'),
    ('Tom-tat-tung-bai/VN016_BaoQuyen_2025_YHCD.docx', 'TOM TAT VN016'),
    ('03_Ban-dich/VN018_AnGiang_2025_YHVN.docx', 'BAN DICH VN018'),
    ('Tom-tat-tung-bai/VN018_AnGiang_2025_YHVN.docx', 'TOM TAT VN018'),
    ('03_Ban-dich/QT005_Alharbi_et_al_2019_SaudiArabia.docx', 'BAN DICH Alharbi'),
    ('Tom-tat-tung-bai/QT005_Alharbi_et_al_2019_SaudiArabia.docx', 'TOM TAT Alharbi'),
    ('03_Ban-dich/QT006_Nakie_et_al_2022_Ethiopia.docx', 'BAN DICH Nakie'),
    ('Tom-tat-tung-bai/QT006_Nakie_et_al_2022_Ethiopia.docx', 'TOM TAT Nakie'),
]

# Patterns to flag - things that may be wrong
SUSPICIOUS = [
    # VN016 Bao Quyen errors
    '1.252', '1252', '38,5%', '38.5%',
    # VN017 Lam errors
    '1.024', '1024', 'bán đô thị', '25,8%', '25.8%',
    # VN018 An Giang errors
    'Lê Minh T.', '600 học sinh', '33,2%', '33.2%',
    # Alharbi
    '366 học sinh', '63,1%', '63.1%',
    # Nakie
    '643 học sinh', '38,9%', '38.9%',
    # Jefferies
    'Canada', '41% nhất', 'cao nhất 41',
    # Term
    'lo âu tổng quát', 'rối loạn lo âu tổng quát',
]

# Correct values for cross-check
CORRECT = [
    # VN016 Bao Quyen 2025 - 501 hs, 86.2%
    '501', '86,2%',
    # VN017 Lam 2022 - 482 hs, 49.0%, bán nông thôn
    '482', '49,0%', 'bán nông thôn',
    # VN018 An Giang 2025 - 366 hs, 61.2%, Nguyễn Đăng Khoa
    '61,2%', 'Nguyễn Đăng Khoa',
    # Alharbi - 1245 hs, 63.4%
    '1.245', '1245', '63,4%',
    # Nakie - 849 hs, 66.7%
    '849', '66,7%',
    # Jefferies - US highest 57.6%
    '57,6%', 'Brazil',
    # Term
    'lo âu lan tỏa',
]

for rel, label in FILES_TO_CHECK:
    full = os.path.join(ROOT, rel)
    if not os.path.exists(full):
        print(f"\n!!! KHONG TIM THAY: {full}")
        continue
    try:
        d = Document(full)
    except Exception as e:
        print(f"\n!!! LOI MO {label}: {e}")
        continue
    print(f"\n========================================")
    print(f"=== {label} ===")
    print(f"========================================")
    print(f"Paragraphs: {len(d.paragraphs)}, Tables: {len(d.tables)}")

    print(f"\n--- LOI TIM THAY (suspicious) ---")
    found_susp = False
    for kw in SUSPICIOUS:
        b = sum(1 for p in d.paragraphs if kw in p.text)
        t = sum(1 for tb in d.tables for row in tb.rows for cell in row.cells if kw in cell.text)
        if b + t > 0:
            print(f"  ✗ '{kw}' body={b} tables={t}")
            found_susp = True
    if not found_susp:
        print(f"  ✓ Khong co loi nghi")

    print(f"\n--- SO LIEU DUNG (correct) ---")
    for kw in CORRECT:
        b = sum(1 for p in d.paragraphs if kw in p.text)
        t = sum(1 for tb in d.tables for row in tb.rows for cell in row.cells if kw in cell.text)
        if b + t > 0:
            print(f"  ✓ '{kw}' body={b} tables={t}")
