# -*- coding: utf-8 -*-
"""Check Saikia context in LA - full para text."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_2_FixSoLieu_27052026.docx')
d = Document(LA)

# Print full text of paragraphs containing Saikia or Sakia
for i, p in enumerate(d.paragraphs):
    if 'Saikia' in p.text or 'Sakia' in p.text:
        print(f"\n--- para {i} ---")
        print(p.text)

# Also search for "288" near Saikia
print("\n\n=== Search '288' ===")
for i, p in enumerate(d.paragraphs):
    if '288' in p.text:
        print(f"para {i}: {p.text[:400]}")
