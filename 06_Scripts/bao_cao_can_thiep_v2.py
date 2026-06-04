# -*- coding: utf-8 -*-
"""
BÁO CÁO CAN THIỆP TÂM LÝ VTN RLLA — phiên bản 2 (chất lượng cao)
- Citation đầy đủ: bài + trang
- 8 biểu đồ matplotlib
- Bảng vừa khổ giấy A4 (lề 3-2cm = nội dung ~16cm)
- Tiếng Việt có dấu đầy đủ
- Phong cách CTH v5
"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
CHARTS = os.path.join(os.path.dirname(__file__), 'charts')
OUT = os.path.join(ROOT, 'Bao cao Can thiep tam ly RLLA VTN - 10042026 v2.docx')

# A4 portrait, lề 3 (T) - 2 (B/L/R) cm → nội dung W = 21 - 5 = 16 cm
PAGE_W = 16.0  # cm

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2)

# ============== HELPERS ==============
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    # remove existing tcW first
    for existing in tcW.findall(qn('w:tcW')):
        tcW.remove(existing)
    tcW.append(w)

def set_grid_cols(t, widths_cm):
    """Set tblGrid widths so Word respects them"""
    tblPr = t._tbl.find(qn('w:tblPr'))
    # Disable autofit on the table
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        # remove old layout if exists
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblPr.append(layout)
    # Replace tblGrid
    tblGrid = t._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        t._tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w_cm in widths_cm:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w_cm * 567)))
        tblGrid.append(gc)
    t._tbl.insert(list(t._tbl).index(t._tbl.findall('.//' + qn('w:tr'))[0]), tblGrid)

def H(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(text, bold=False, italic=False, size=12, color=None, align=None):
    p = doc.add_paragraph()
    if align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb = color
    return p

def Pmix(parts):
    """parts = [(text, {bold, italic, color, size}), ...]"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 12))
        r.bold = opts.get('bold', False)
        r.italic = opts.get('italic', False)
        if 'color' in opts: r.font.color.rgb = opts['color']
    return p

def cite(text):
    """Format citation italic gray"""
    return (text, {'italic': True, 'size': 10.5, 'color': RGBColor(0x55, 0x55, 0x55)})

def table(headers, rows, widths, header_color='D9E2F3', font_size=10):
    """Bảng có viền, cột co theo cm. Tổng widths phải <= PAGE_W"""
    assert sum(widths) <= PAGE_W + 0.05, f'Tổng cột {sum(widths)}cm vượt {PAGE_W}cm'
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    t.allow_autofit = False
    set_grid_cols(t, widths)
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(font_size)
        shade(c, header_color)
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(font_size)
    return t

def img(filename, width_cm=15.0, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(CHARTS, filename)
    p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.italic=True

# ============================================================
# TRANG TIÊU ĐỀ
# ============================================================
H('BÁO CÁO TỔNG HỢP', level=1)
H('BIỆN PHÁP HỖ TRỢ TÂM LÝ VÀ CAN THIỆP TÂM LÝ\nCHO TRẺ VỊ THÀNH NIÊN BỊ RỐI LOẠN LO ÂU', level=1)
P('Tổng hợp 15 công trình nghiên cứu — Phân theo 3 vùng địa lý:\n'
  '(1) Việt Nam | (2) Châu Á (ngoài Việt Nam) | (3) Châu Âu — Châu Úc — Châu Mỹ',
  italic=True, align='center')
P('Cập nhật: 10/04/2026 | Phong cách: Công Thị Hằng v5\n'
  'Dữ liệu trích từ hệ thống 58 bài nghiên cứu của dự án Lo âu',
  italic=True, align='center', size=11)

# ============================================================
# TÓM LƯỢC ĐIỀU HÀNH
# ============================================================
H('TÓM LƯỢC ĐIỀU HÀNH', level=2)

P('Báo cáo này tổng hợp 15 công trình nghiên cứu về biện pháp hỗ trợ và can thiệp tâm lý '
  'cho rối loạn lo âu (RLLA) ở trẻ vị thành niên (VTN), được phân theo ba vùng địa lý nhằm '
  'làm rõ bằng chứng khu vực, mức độ phù hợp văn hoá và khả năng áp dụng cho Việt Nam. '
  'Phương pháp tổng hợp kết hợp đối chiếu liên bài, phân tích phản biện và xếp hạng bằng chứng '
  'theo chất lượng thiết kế nghiên cứu.', align='justify')

P('Sáu phát hiện then chốt:', bold=True)
P('Thứ nhất, Liệu pháp Nhận thức – Hành vi (CBT) là phương pháp được chứng minh hiệu quả nhất '
  'và nhất quán xuyên suốt cả ba vùng địa lý. Bằng chứng vàng từ thử nghiệm CAMS '
  '(Walkup et al. 2008, NEJM 359:2753–66) cho thấy CBT kết hợp Sertraline đạt 80,7 % đáp ứng '
  'so với 59,7 % cho CBT đơn thuần và 23,7 % cho giả dược (n = 488 trẻ 7–17 tuổi).', align='justify')
P('Thứ hai, CBT qua Internet (iCBT) cho rối loạn lo âu xã hội (SAD) đứng hạng 1 trong phân tích '
  'tổng hợp mạng 30 RCT (Xian, Zhang & Jiang 2024, J Affect Disord 365:614–627; SUCRA 71,2 %), '
  'và kích thước hiệu ứng đặc biệt LỚN khi can thiệp được thiết kế riêng cho SAD '
  '(Walder et al. 2025, JMIR; Hedges g = 0,878).', align='justify')
P('Thứ ba, ứng dụng di động CBT (Mobile CBT) cho hiệu quả MẠNH với trầm cảm (87,5 % nghiên cứu '
  'dương tính) nhưng YẾU với lo âu khi không thiết kế riêng (Qiaochu & Wang 2025, Clin Psychol '
  'Psychother 32:e70173). Tuy nhiên RCT ứng dụng Maya của Bress et al. 2024 cho thấy hiệu lực '
  'rất lớn (Cohen d = 1,04 sau 12 tuần, JAMA Network Open 7(8):e2428372).', align='justify')
P('Thứ tư, can thiệp PHỔ QUÁT tại trường (universal CBT) chỉ cho hiệu quả nhỏ và chất lượng '
  'bằng chứng nền thấp (Zhang, Liang & Kang 2026, J Clin Psychol 82:248–259, Bayesian MA 31 RCT, '
  'n = 19.865). Mô hình CHỈ ĐỊNH (targeted) cho học sinh có triệu chứng — như BESST UK '
  '(Brown & Carter 2025, J Mental Health 34(4):357–361) — hứa hẹn hơn và phù hợp với cảnh báo '
  'của Brown & Carter về thất bại Mindfulness phổ quát ở 8.376 học sinh UK.', align='justify')
P('Thứ năm, KHOẢNG TRỐNG NGHIÊN CỨU TẠI VIỆT NAM CỰC LỚN: trong hệ thống 58 bài nghiên cứu '
  'không có một thử nghiệm lâm sàng ngẫu nhiên (RCT) nào về can thiệp tâm lý cho VTN có RLLA '
  'tại Việt Nam. Công trình duy nhất từ tác giả Việt Nam là luận án Tiến sĩ Y học của '
  'Trần Nguyễn Ngọc (2018, Đại học Y Hà Nội — Viện Sức khoẻ Tâm thần Bạch Mai), nhưng đối tượng '
  'là người lớn nội trú, không phải VTN.', align='justify')
P('Thứ sáu, khu vực Đông Á và Thái Bình Dương ở các nước có thu nhập trung bình – thấp (LMIC) '
  'còn thiếu nghiêm trọng các can thiệp ở cấp CỘNG ĐỒNG, GIA ĐÌNH và dịch vụ ĐÁP ỨNG dài hạn '
  '(Menon et al. 2025, Asia Pac J Public Health 37(4):332–346, scoping review 69 nghiên cứu, '
  '12 quốc gia). Phát hiện này định hướng mạnh cho thiết kế đề cương Việt Nam.', align='justify')

# Chart 8: 3 regions overview
img('chart8_3regions.png', width_cm=15.5,
    caption='Biểu đồ 8. Phân bố nghiên cứu can thiệp tâm lý cho RLLA theo vùng địa lý '
            '(tổng hợp 15 bài trong hệ thống 58 bài, dự án Lo âu, 04/2026).')

# ============================================================
# PHẦN I — VIỆT NAM
# ============================================================
doc.add_page_break()
H('PHẦN I — VIỆT NAM', level=1)

P('Trong hệ thống 58 bài nghiên cứu của dự án, chỉ có 1 công trình duy nhất từ tác giả Việt Nam '
  'thuộc nhóm chủ đề CAN THIỆP tâm lý cho rối loạn lo âu — luận án Tiến sĩ Y học của '
  'Trần Nguyễn Ngọc (2018). Tuy nhiên, đối tượng nghiên cứu là bệnh nhân NGƯỜI LỚN nội trú '
  'tại Viện Sức khoẻ Tâm thần Bạch Mai, không phải vị thành niên. Đây là khoảng trống nghiên '
  'cứu cực lớn và là chỗ trống mà đề cương của chúng ta hướng đến lấp đầy.', align='justify', italic=True)

H('1.1. Trần Nguyễn Ngọc (2018) — Liệu pháp Thư giãn – Luyện tập cho RLLA lan toả', level=2)

H('1.1.1. Thông tin thư mục', level=3)
table(
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Trần Nguyễn Ngọc'],
        ['Hướng dẫn khoa học', 'PGS. TS. Nguyễn Kim Việt'],
        ['Loại tài liệu', 'Luận án Tiến sĩ Y học, chuyên ngành Tâm thần (mã số 62720148)'],
        ['Đơn vị', 'Trường Đại học Y Hà Nội — Viện Sức khoẻ Tâm thần, BV Bạch Mai'],
        ['Năm bảo vệ', '2018'],
        ['Số trang', '177 trang'],
        ['Đối tượng', '170 bệnh nhân RLLA lan toả (mã F41.1, ICD-10), nội trú Viện SKTT BV Bạch Mai'],
        ['Đặc điểm mẫu', '65 nam (38,2 %) + 105 nữ (61,8 %); tuổi trung bình 43,2 ± 13,6 (Bảng 3.1, tr. 57)'],
        ['Lưu ý quan trọng', 'NGƯỜI LỚN — không phải VTN. Không thể ngoại suy trực tiếp cho học sinh.'],
        ['Can thiệp', 'Liệu pháp Thư giãn – Luyện tập, 20 buổi trong 4 tuần (60 phút/buổi, 5–7 BN/nhóm)'],
        ['Thiết kế', 'Mô tả cắt ngang (n = 170) + can thiệp lâm sàng trước–sau (n = 99)'],
        ['Nguyên tắc', 'KHÔNG kết hợp thuốc — đánh giá thuần liệu pháp tâm lý – thể chất'],
        ['Công cụ đánh giá', 'HAM-A (Hamilton Anxiety) + CGI + PSQI (giấc ngủ) + EPI (nhân cách Eysenck)'],
        ['Thời điểm đo', 'T0 (vào viện), T2 (tuần 2), T4 (tuần 4 — kết thúc)'],
    ],
    widths=[4.0, 11.5]
)

H('1.1.2. Mô tả can thiệp', level=3)
P('Liệu pháp Thư giãn – Luyện tập của Viện Sức khoẻ Tâm thần Bạch Mai gồm ba thành phần cốt lõi '
  'được trình bày trong Chương 1 của luận án (tr. 26–41):', align='justify')
P('• Phần TẬP THƯ GIÃN: kỹ thuật thư giãn cơ tiến triển (Jacobson) và thư giãn tự sinh (Schultz). '
  'Mục tiêu kích hoạt hệ thần kinh đối giao cảm, giảm trương lực cơ và hạ phản ứng stress.', align='justify')
P('• Phần TẬP THỞ: thở bụng, thở chậm sâu, thở luân phiên theo nguyên tắc Yoga. Tác động lên '
  'biến thiên tần số tim (HRV), cân bằng hệ thần kinh tự chủ (Trần Nguyễn Ngọc 2018, tr. 38–41).', align='justify')
P('• Phần TẬP YOGA: các tư thế (asana) cơ bản phù hợp bệnh nhân tâm thần — không vận động cường '
  'độ cao. Tác giả dẫn nhiều nghiên cứu cho thấy Yoga cải thiện sức khoẻ thể chất, độ dẻo cơ thể, '
  'giảm cortisol và rối loạn mỡ máu (tr. 39–40).', align='justify')

P('Cấu trúc một buổi 60 phút chia làm 5 phần (tr. 45):', italic=True, align='justify')
table(
    ['Phần', 'Thời gian', 'Nội dung', 'Người phụ trách'],
    [
        ['1', '15 phút', 'Đánh giá kết quả buổi trước; thảo luận; đo sinh tồn', 'Bác sĩ + điều dưỡng'],
        ['2', '10 phút', 'Khởi động cơ thể', 'Cử nhân tâm lý'],
        ['3', '20 phút', 'Tập chính: thư giãn + thở + Yoga', 'Cử nhân tâm lý'],
        ['4', '10 phút', 'Tập trung tâm trí', 'Cử nhân tâm lý'],
        ['5', '5 phút', 'Kết thúc, đo sinh tồn lại', 'Điều dưỡng'],
    ],
    widths=[1.2, 2.0, 8.5, 3.8]
)
P('Mỗi buổi tập có ít nhất 1 bác sĩ + 1 cán bộ tâm lý + 1 điều dưỡng theo dõi. Tổng cộng 20 '
  'buổi trong 4 tuần. Bệnh nhân tập tại Phòng Tâm lý lâm sàng của Viện Sức khoẻ Tâm thần.',
  italic=True, align='justify')

H('1.1.3. Kết quả chính', level=3)
P('Sau 4 tuần điều trị (n = 99 bệnh nhân hoàn thành), kết quả trên thang HAM-A cho thấy sự '
  'thuyên giảm rõ rệt và có ý nghĩa thống kê mạnh (Trần Nguyễn Ngọc 2018, Bảng 3.20, tr. 76):',
  align='justify')

table(
    ['Mức độ lo âu', 'T0 (vào viện)', 'T2 (tuần 2)', 'T4 (tuần 4)', 'p (T0–T2)', 'p (T0–T4)'],
    [
        ['Nhẹ', '34,3 %', '53,6 %', '52,5 %', '< 0,0001', '0,0001'],
        ['Vừa', '20,2 %', '24,2 %', '36,4 %', '0,1582', '< 0,0001'],
        ['Nặng', '45,5 %', '22,2 %', '11,1 %', '< 0,0001', '< 0,0001'],
    ],
    widths=[3.5, 2.5, 2.5, 2.5, 2.5, 2.5]
)
P('Bảng 1. Mức độ lo âu HAM-A theo thời gian điều trị (n = 99). '
  'Nguồn: Trần Nguyễn Ngọc 2018, Bảng 3.20, tr. 76.', italic=True, size=10)

img('chart1_HAMA_TranNguyenNgoc.png', width_cm=15.5)

P('Đáng chú ý, tỷ lệ bệnh nhân ở mức NẶNG giảm mạnh từ 45,5 % (T0) xuống còn 11,1 % (T4) — '
  'giảm gần 4 lần. Tỷ lệ mức VỪA tăng tạm thời ở T2 (24,2 %) và rõ rệt ở T4 (36,4 %), phản ánh '
  'hiện tượng các bệnh nhân nặng "trượt xuống" mức vừa rồi tiếp tục cải thiện thành mức nhẹ. '
  'Đây là kiểu đáp ứng điển hình của can thiệp tâm lý dài hạn.', align='justify')

P('Về 22 triệu chứng riêng lẻ (Bảng 3.22 tr. 77), tổng số triệu chứng trung bình giảm từ '
  '11,8 ± 3,5 (T0) xuống 5,1 ± 4,9 (T4), với p < 0,0001. Nhóm triệu chứng kích thích thần kinh '
  'thực vật cải thiện đặc biệt mạnh:', align='justify')

table(
    ['Triệu chứng', 'T0', 'T2', 'T4', 'p (T0–T4)'],
    [
        ['Hồi hộp / tim đập nhanh', '88,9 %', '73,7 %', '43,4 %', '< 0,0001'],
        ['Vã mồ hôi', '59,6 %', '36,3 %', '16,1 %', '< 0,0001'],
        ['Run', '57,6 %', '34,3 %', '17,1 %', '< 0,0001'],
        ['Khô miệng', '38,4 %', '25,2 %', '16,2 %', '< 0,0001'],
        ['Khó thở', '56,6 %', '46,4 %', '25,2 %', '< 0,0001'],
        ['Chóng mặt / không vững', '66,7 %', '48,4 %', '32,3 %', '< 0,0001'],
        ['Sợ mất kiềm chế', '31,3 %', '20,2 %', '10,1 %', '< 0,0001'],
        ['Sợ bị chết', '33,3 %', '19,1 %', '9,1 %', '< 0,0001'],
        ['Dễ giật mình', '61,6 %', '28,2 %', '13,1 %', '< 0,0001'],
        ['Khó tập trung', '65,7 %', '29,2 %', '11,1 %', '< 0,0001'],
    ],
    widths=[6.0, 2.0, 2.0, 2.0, 3.0]
)
P('Bảng 2. Thuyên giảm 10 triệu chứng tiêu biểu sau 4 tuần điều trị (n = 99). '
  'Nguồn: Trần Nguyễn Ngọc 2018, Bảng 3.23 tr. 78, Bảng 3.26 tr. 80, Bảng 3.28 tr. 81.',
  italic=True, size=10)

img('chart2_TKTV_TranNguyenNgoc.png', width_cm=15.5)

P('Trên thang đánh giá tổng thể bệnh trạng CGI (Bảng 3.29 tr. 82), tỷ lệ bệnh nhân ở mức '
  'NẶNG / RẤT NẶNG giảm từ 52,5 % (T0) xuống 13,1 % (T2) và còn 9,1 % (T4) — giảm gần 6 lần. '
  'Cùng lúc, tỷ lệ ở mức NHẸ tăng từ 0 % (T0) lên 17,1 % (T2) và 59,6 % (T4). Tất cả các so '
  'sánh đều có p < 0,0001.', align='justify')

H('1.1.4. Đặc điểm lâm sàng đáng chú ý (n = 170)', level=3)
P('Phân tích mẫu mô tả của Trần Nguyễn Ngọc (2018, tr. 62–75) cho thấy nhiều đặc điểm có ý '
  'nghĩa thực hành lâm sàng:', align='justify')
P('• 45,3 % bệnh nhân có sang chấn tâm lý — trong đó 32,4 % là sang chấn trường diễn '
  '(Bảng 3.8, tr. 65). Sang chấn gia đình chiếm 33,5 %, tai nạn / bệnh tật 27,1 %.', align='justify')
P('• 12,4 % bệnh nhân có TRẦM CẢM kèm theo, với chênh lệch giới rõ rệt (nữ 15,2 % so với nam '
  '7,7 %, Bảng 3.9 tr. 66).', align='justify')
P('• Sử dụng chất cộng đồng nam – nữ rất khác biệt: nam dùng thuốc lá 67,7 % và rượu 44,6 %, '
  'trong khi nữ chỉ 1,9 % cho cả hai (Bảng 3.9 tr. 66).', align='justify')
P('• Theo Montgomery (2010, được Trần Nguyễn Ngọc 2018 trích dẫn ở tr. 1), chỉ khoảng 13,3 % '
  'bệnh nhân RLLA đến khám VÌ triệu chứng lo âu — phần lớn đến vì các triệu chứng cơ thể, '
  'dẫn đến chẩn đoán muộn và chuyển khoa nhiều lần.', align='justify')

H('1.1.5. Quan điểm phản biện', level=3)
P('Điểm mạnh:', bold=True)
P('• Đây là LUẬN ÁN ĐẦU TIÊN tại Việt Nam đánh giá hiệu quả liệu pháp Thư giãn – Luyện tập cho '
  'RLLA lan toả (lời cam đoan của tác giả ở tr. 3: "Công trình này không trùng lặp với bất kỳ '
  'nghiên cứu nào khác đã được công bố tại Việt Nam").', align='justify')
P('• Cỡ mẫu lớn cho mô tả lâm sàng (n = 170) và đủ cho can thiệp trước–sau (n = 99), tính theo '
  'công thức dựa trên Naomi Breslau với độ chính xác Δ = 6,5 và α = 0,05 (tr. 43).', align='justify')
P('• Quy trình chuẩn 20 buổi có giám sát đa chuyên ngành (bác sĩ, cán bộ tâm lý, điều dưỡng), '
  'thực hiện tại cơ sở chuyên khoa, đảm bảo tuân thủ.', align='justify')
P('• Hiệu quả có ý nghĩa thống kê mạnh (p < 0,0001) trên CẢ HAM-A tổng và 22 triệu chứng riêng lẻ.', align='justify')
P('• Loại trừ thuốc — đánh giá thuần liệu pháp tâm lý – thể chất, giúp xác định hiệu lực '
  '"net" của bản thân liệu pháp và tăng khả năng nhân rộng tại các cơ sở y tế tuyến cơ sở.', align='justify')
P('• Kỹ thuật rẻ tiền, đơn giản, dễ thực hiện — phù hợp triển khai tại trường học, trạm y tế '
  'xã, hoặc cộng đồng.', align='justify')

P('Hạn chế nghiêm trọng cho mục đích áp dụng cho VTN:', bold=True, color=RGBColor(0xCC,0,0))
P('• ĐỐI TƯỢNG NGƯỜI LỚN — tuổi trung bình 43,2 ± 13,6 (Bảng 3.1 tr. 57). Không thể ngoại suy '
  'trực tiếp cho học sinh THCS / THPT do khác biệt phát triển não bộ, môi trường tâm lý xã hội, '
  'và kỹ năng nhận thức – cảm xúc.', align='justify')
P('• THIẾT KẾ TRƯỚC – SAU KHÔNG ĐỐI CHỨNG — không phải RCT. Không thể loại trừ hiệu ứng giả '
  'dược, hồi quy về trung bình, và hiệu ứng Hawthorne (cải thiện do được chú ý chứ không phải '
  'do liệu pháp).', align='justify')
P('• Mẫu chỉ tại Bệnh viện Bạch Mai (nội trú) — chọn lọc mạnh: bệnh nhân thường nặng, có động '
  'lực điều trị cao, tuân thủ tốt. Không đại diện bệnh nhân RLLA cộng đồng.', align='justify')
P('• Theo dõi chỉ 4 tuần — không đánh giá hiệu quả dài hạn 3 tháng, 6 tháng, 1 năm. So với '
  'chuẩn hiện đại (như Zhang, Liang & Kang 2026 đo tới 1 năm), thời gian theo dõi quá ngắn.', align='justify')
P('• Không có nhóm đối chứng (chờ đợi / thuốc / CBT) — không thể so sánh hiệu lực với các '
  'phương pháp khác.', align='justify')
P('• Tỷ lệ rút khỏi nghiên cứu 41,8 % (170 → 99) — RẤT CAO. Có thể có thiên lệch người sống '
  'sót (survivor bias) vì chỉ những bệnh nhân đáp ứng tốt mới hoàn thành đủ 20 buổi.', align='justify')

H('1.1.6. Áp dụng cho đề cương VTN Việt Nam', level=3)
P('Mặc dù đối tượng không phải VTN, luận án Trần Nguyễn Ngọc (2018) cung cấp ba bài học quý '
  'cho thiết kế can thiệp Việt Nam:', bold=True, align='justify')
P('Một là, liệu pháp Thư giãn – Luyện tập có thể là MODULE BỔ TRỢ cho can thiệp CBT trường '
  'học (giai đoạn 2 đề cương) — phù hợp với Li et al. (2025, BMC Psychiatry 25:809) cho thấy '
  'PE (hoạt động thể chất) đứng SUCRA 0,51 (Bảng kết quả NMA tr. 6) và Cao et al. (2025, '
  'Frontiers in Psychiatry 16:1594658) cho thấy can thiệp resilience tại trường có hiệu quả nhỏ '
  'đến trung bình.', align='justify')
P('Hai là, cần ADAPTATION cho VTN: thời gian buổi ngắn hơn (30–45 phút), nội dung trẻ hơn '
  '(âm nhạc, trò chơi), kết hợp peer learning để tăng engagement. Cần thiết kế RCT có đối '
  'chứng để cung cấp bằng chứng vững chắc.', align='justify')
P('Ba là, KHOẢNG TRỐNG SỐ 1 của lĩnh vực được xác định rõ: Việt Nam cần luận án / RCT đầu tiên '
  'về can thiệp tâm lý cho VTN có RLLA — đối tượng học sinh THCS / THPT chưa được nghiên cứu '
  'can thiệp. Đây chính là trọng tâm của đề cương dự án.', align='justify')

# ============================================================
# PHẦN II — CHÂU Á NGOÀI VN
# ============================================================
doc.add_page_break()
H('PHẦN II — CHÂU Á (ngoài Việt Nam)', level=1)

P('Phần này tổng hợp 7 công trình từ Trung Quốc, Nhật Bản, Sri Lanka, các nước Đông Nam Á, '
  'và scoping review LMIC khu vực Đông Á – Thái Bình Dương. Đây là vùng có tính tương đồng văn '
  'hoá cao nhất với Việt Nam (Á Đông, gia đình mở rộng, áp lực học tập, kỳ thị bệnh tâm thần), '
  'do đó có giá trị tham chiếu trực tiếp cho thiết kế đề cương Việt Nam.', align='justify', italic=True)

H('2.1. Tổng quan 7 nghiên cứu can thiệp Châu Á', level=2)
table(
    ['#', 'Tác giả / Năm', 'Quốc gia', 'Loại nghiên cứu', 'Đối tượng / Cỡ mẫu', 'Phát hiện then chốt'],
    [
        ['41', 'Praptomojati & Hartanto 2024', 'Tổng hợp ĐNA', 'SR 7 NC', 'Trẻ + VTN', 'CA-CBT thích ứng văn hoá; 0/7 NC từ VN'],
        ['42', 'De Silva et al. 2024', 'Sri Lanka', 'Cluster RCT', '720 HS lớp 9 / 36 trường', 'CBT do GV dạy; β = −0,096; p = 0,038'],
        ['43', 'Xian, Zhang & Jiang 2024', 'Trung Quốc', 'NMA Bayesian', '30 RCT, 1.547 VTN SAD', 'iCBT SUCRA 71,2 % hạng 1; gCBT 68,4 %'],
        ['51', 'Sasaki et al. 2024 (B11)', 'Nhật Bản', 'RCT đa trung tâm', '77 HS-SV (38 + 39) tại 6 ĐH + 1 THPT', 'iCBT tự học; OR đáp ứng = 4,97; OR phục hồi = 3,95'],
        ['57*', 'Qiaochu & Wang 2025', 'Trung Quốc', 'SR 9 RCT', '2.479 trẻ + người trẻ', 'Mobile CBT: trầm cảm 7/8; lo âu chỉ 2/6'],
        ['58*', 'Menon et al. 2025', 'LMIC ĐÁ-TBD', 'Scoping 69 NC', '12 quốc gia', 'Thiếu cộng đồng, gia đình, dài hạn'],
        ['29', 'Li et al. 2025 (BMC NMA)', 'Trung Quốc', 'NMA Bayesian', '30 RCT, 1.711 trẻ', 'CBT SUCRA 0,66 (16 RCT); PE 0,51 (2 RCT)'],
    ],
    widths=[0.8, 3.0, 1.7, 2.3, 3.2, 4.5]
)
P('Bảng 3. Tổng quan 7 nghiên cứu can thiệp Châu Á (ngoài Việt Nam). '
  '* Các bài đánh dấu * chỉ có thông tin abstract (paywall).', italic=True, size=10)

H('2.2. Praptomojati & Hartanto (2024) — CA-CBT cho rối loạn lo âu ở Đông Nam Á', level=2)
P('Tổng quan hệ thống về Liệu pháp Nhận thức – Hành vi Thích ứng Văn hoá (Culturally Adapted '
  'CBT — CA-CBT) cho rối loạn lo âu tại Đông Nam Á, đăng trên Asian Journal of Psychiatry '
  '(Q1, IF ≈ 12 — tạp chí tâm thần châu Á uy tín nhất). Tìm kiếm hệ thống PubMed, PsycINFO, '
  'Embase, CENTRAL, GARUDA và Google Scholar đã chọn được 7 nghiên cứu (1 RCT, 3 bán thực '
  'nghiệm, 3 ca lâm sàng) trên 11 nước ĐNA — trong đó có 0 nghiên cứu nào từ Việt Nam '
  '(Praptomojati & Hartanto 2024, A J Psychiatry vol. 92, tr. 1–9).', align='justify')

P('Hai trong số 7 nghiên cứu thực hiện thích ứng đa thành phần, hai nghiên cứu sửa đổi thành '
  'phần CỐT LÕI bằng tích hợp tôn giáo (Hồi giáo ở Indonesia, Phật giáo ở Thái Lan) hoặc ngôn '
  'ngữ địa phương. Khung CTAF (Cultural Treatment Adaptation Framework) phân loại rõ thích '
  'ứng ngoại vi (peripheral — ngôn ngữ, hình ảnh) so với thích ứng cốt lõi (deep — niềm tin, '
  'hệ giá trị).', align='justify')

P('Phản biện: SR khu vực ĐNA RẤT QUÝ cho Việt Nam vì xác định rõ khoảng trống — Việt Nam là '
  'một trong những nước có 0 nghiên cứu CA-CBT trong khu vực. Áp dụng cho dự án: cần phát '
  'triển CA-CBT phiên bản tiếng Việt với thích ứng văn hoá Á Đông (hệ giá trị Khổng giáo, gia '
  'đình mở rộng) và yếu tố tôn giáo (Phật giáo, Công giáo). Phù hợp với phát hiện của '
  'Dong, Wang & Lin (2025, PLOS ONE 20(9):e0328785, tr. 8) về "fear of letting down others" '
  'ở 60,3 % học sinh Trung Quốc — một cấu trúc tâm lý đặc thù Á Đông cần can thiệp riêng.',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('2.3. De Silva et al. (2024) — Cluster RCT can thiệp CBT trường tại Sri Lanka', level=2)
P('Thử nghiệm ngẫu nhiên có đối chứng theo cụm (cluster RCT) — phân ngẫu nhiên 36 TRƯỜNG '
  'học (không phải học sinh): 18 trường can thiệp + 18 trường đối chứng. Mỗi nhánh 360 học '
  'sinh lớp 9 (~ 14 tuổi), tổng cộng 720 học sinh. Đăng trên Child and Adolescent Psychiatry '
  'and Mental Health vol. 18, article 108, tr. 1–12, DOI 10.1186/s13034-024-00799-9 '
  '(De Silva et al. 2024).', align='justify')

P('Can thiệp gồm 8 phiên CBT mỗi tuần, 40 phút/phiên, do GIÁO VIÊN cung cấp sau khi được '
  'đào tạo. Chương trình bao gồm tâm lý giáo dục, kỹ thuật thở, tái cấu trúc nhận thức, kỹ '
  'năng giải quyết vấn đề và kỹ năng đối phó. Theo dõi 3 tháng sau can thiệp.', align='justify')

P('Kết quả chính (De Silva et al. 2024, Bảng 3 tr. 7):', bold=True)
P('• Tỷ lệ mất mẫu < 1 % — tuân thủ rất cao, chứng tỏ tính khả thi của mô hình.', align='justify')
P('• Lo âu giảm có ý nghĩa thống kê sau 3 tháng: β = −0,096 (KTC 95 %: −0,186 đến −0,005, '
  'p = 0,038).', align='justify')
P('• Tự trọng (self-esteem) tăng có ý nghĩa: β = +0,811 sau can thiệp (cùng bảng).', align='justify')

P('Phản biện: Sri Lanka là quốc gia LMIC Nam Á — bối cảnh nguồn lực hạn chế tương tự Việt Nam. '
  'Mô hình "GV cung cấp CBT" rẻ tiền, có thể nhân rộng. Đây là MÔ HÌNH KHẢ THI nhất cho Việt '
  'Nam trong số tất cả các nghiên cứu được tổng hợp. Tuy có khác biệt với Brown & Carter '
  '(2025, J Mental Health 34(4):358) — họ cho rằng "chuyên gia lâm sàng thường hiệu quả hơn '
  'giáo viên" — nhưng Sri Lanka cho thấy GV vẫn ĐỦ với đào tạo phù hợp. Cần lặp lại tại Việt '
  'Nam để có bằng chứng nội địa.', italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('2.4. Xian, Zhang & Jiang (2024) — Network Meta-Analysis 30 RCT cho SAD', level=2)
P('Phân tích tổng hợp mạng Bayesian 30 RCT, n = 1.547 vị thành niên có Lo âu Xã hội (SAD), '
  '9 liệu pháp + 3 nhóm đối chứng. Đăng trên Journal of Affective Disorders vol. 365, '
  'tr. 614–627 (Xian, Zhang & Jiang 2024). PROSPERO CRD42023476829.', align='justify')

P('Đây là phân tích tổng hợp mạng đầu tiên tập trung riêng cho SAD ở vị thành niên — bằng '
  'chứng vàng cho việc xếp hạng các can thiệp tâm lý. Kết quả SUCRA (Bảng 2 tr. 619):', align='justify')

table(
    ['Hạng', 'Liệu pháp', 'SUCRA (%)', 'Ghi chú'],
    [
        ['1', 'iCBT (CBT internet)', '71,2', 'Tự học online + hướng dẫn — phù hợp VTN số hoá'],
        ['2', 'gCBT (CBT nhóm)', '68,4', 'Nhóm 6–12 HS, 8–16 tuần — KHẢ THI nhất cho trường VN'],
        ['3', 'I-CBT (CBT cá nhân)', '66,0', 'Tốn chuyên gia — khó cho LMIC'],
        ['4', 'SET-C (Training XH)', '~ 60', 'Kỹ năng XH + phơi nhiễm — kết hợp được'],
        ['5', 'IPT (Liệu pháp giữa cá nhân)', '~ 55', 'Ít nghiên cứu'],
        ['6', 'AT (Huấn luyện chú ý)', '~ 50', 'Còn ở giai đoạn thực nghiệm'],
        ['—', 'WL / NT / PBO (đối chứng)', '< 20', 'Đối chứng — không hiệu quả'],
    ],
    widths=[1.0, 5.0, 2.5, 7.0]
)
P('Bảng 4. Xếp hạng can thiệp cho rối loạn lo âu xã hội ở vị thành niên. '
  'Nguồn: Xian, Zhang & Jiang 2024, Bảng 2, tr. 619.', italic=True, size=10)

img('chart3_NMA_SAD_Xian.png', width_cm=15.5)

P('Hai phát hiện phụ quan trọng (Bảng 3 tr. 620–621): iCBT cho hiệu quả tốt nhất với TRẦM CẢM '
  'kèm theo (SUCRA 92,2 %), trong khi gCBT xếp hạng 1 cho cải thiện CHỨC NĂNG xã hội '
  '(SUCRA 89,6 %). Như vậy, kết hợp iCBT (giảm triệu chứng) với gCBT (phục hồi chức năng) có '
  'thể là chiến lược tối ưu.', align='justify')

P('Phản biện: NMA cho phép so sánh GIÁN TIẾP tất cả can thiệp — bằng chứng MẠNH NHẤT về xếp '
  'hạng có thể thu được không cần thử nghiệm head-to-head. iCBT hạng 1 cho SAD — nên ưu tiên '
  'phát triển app tiếng Việt. gCBT hạng 2 nhưng có ưu thế đặc biệt cho phục hồi chức năng — '
  'phù hợp triển khai ở trường. Hạn chế: chỉ tập trung SAD (không phải GAD); 30 RCT chủ yếu '
  'từ phương Tây và Đông Á phát triển (Trung Quốc, Hàn Quốc, Nhật Bản).',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('2.5. Sasaki et al. (2024) — RCT iCBT đa trung tâm tại Nhật Bản', level=2)
P('Thử nghiệm ngẫu nhiên đa trung tâm tại 6 đại học và 1 trường THPT Nhật Bản, từ 12/2022 '
  'đến 10/2023. Đối tượng là học sinh / sinh viên có triệu chứng SAD DƯỚI NGƯỠNG '
  '(subthreshold) — chưa đủ tiêu chuẩn chẩn đoán đầy đủ. n = 77 (38 can thiệp + 39 đối '
  'chứng). Đăng trên JMIR Pediatrics and Parenting vol. 7, tr. 1–13. UMIN000049768. '
  '(Sasaki et al. 2024).', align='justify')

P('Can thiệp: iCBT TỰ HỌC HOÀN TOÀN (không có hướng dẫn người), 8 module qua web. Chi phí '
  'thực hiện thấp nhất trong các mô hình can thiệp được tổng hợp.', align='justify')

P('Kết quả chính (Sasaki et al. 2024, Bảng 2 tr. 7):', bold=True)
P('• Tỷ lệ đáp ứng (response rate) 61 %, OR = 4,97 so với nhóm chờ.', align='justify')
P('• Tỷ lệ phục hồi (recovery rate) 68 %, OR = 3,95.', align='justify')
P('• ANCOVA với trầm cảm là biến đồng: nhóm can thiệp giảm có ý nghĩa thống kê triệu chứng '
  'lo âu xã hội so với nhóm đối chứng (cùng bảng).', align='justify')

P('Phản biện: RCT đa trung tâm — bằng chứng cao. SAD DƯỚI NGƯỠNG là dân số rất quan trọng '
  '(rất lớn về số lượng) — can thiệp ở giai đoạn này có thể NGĂN CHẶN tiến triển thành rối '
  'loạn đầy đủ, là chiến lược "indicated prevention". Mô hình iCBT TỰ HỌC HOÀN TOÀN có chi '
  'phí gần như bằng 0 sau khi phát triển nội dung — phù hợp triển khai đại trà cho Việt Nam '
  'qua web hoặc app. Đây là một trong hai mô hình khả thi nhất (cùng với mô hình GV cung cấp '
  'CBT của Sri Lanka) cho LMIC.', italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('2.6. Qiaochu & Wang (2025) — SR 9 RCT về Mobile CBT', level=2)
P('Tổng quan hệ thống 9 RCT về CBT cung cấp qua ứng dụng smartphone cho trầm cảm và lo âu ở '
  'trẻ em + người trẻ. Tổng N = 2.479. Đăng trên Clinical Psychology and Psychotherapy vol. '
  '32(6), e70173 (Qiaochu & Wang 2025). Bài này hiện chỉ có abstract qua paywall Wiley.', align='justify')

P('Phát hiện then chốt (theo abstract Crossref):', bold=True)
P('• 7 trên 8 nghiên cứu (87,5 %) đo TRẦM CẢM cho thấy giảm có ý nghĩa thống kê.', align='justify')
P('• Chỉ 2 trên 6 nghiên cứu (33,3 %) đo LO ÂU cho thấy hiệu quả có ý nghĩa.', align='justify')
P('• Kết luận: Mobile CBT có hứa hẹn cho trầm cảm ở trẻ em và người trẻ, nhưng bằng chứng '
  'cho lo âu còn HẠN CHẾ.', align='justify')

P('Phản biện: Phát hiện này tưởng chừng mâu thuẫn với Sasaki et al. 2024 (iCBT Nhật cho SAD '
  'dưới ngưỡng — dương tính) và Walder et al. 2025 (DMHI g = 0,878 cho SAD-specific). Lý do '
  'có thể: (1) các app trong Qiaochu & Wang được thiết kế cho TRẦM CẢM tổng hợp, không '
  'SAD-specific; (2) "mobile" và "internet" là hai mô hình khác nhau (smartphone app vs '
  'web-based với hỗ trợ); (3) thiếu hướng dẫn người (guided support). Cần đối chiếu thêm khi '
  'có full PDF. Bài học cho Việt Nam: nếu phát triển app tiếng Việt, nên bắt đầu với module '
  'TRẦM CẢM (bằng chứng mạnh) trước khi mở rộng sang lo âu, hoặc chỉ làm app dành riêng cho '
  'lo âu xã hội với hướng dẫn người.', italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('2.7. Menon et al. (2025) — Scoping Review LMIC Đông Á – Thái Bình Dương', level=2)
P('Tổng quan phạm vi (scoping review) về can thiệp sức khoẻ tâm thần và sức khoẻ tâm lý xã '
  'hội cho trẻ em + vị thành niên tại các nước thu nhập thấp – trung bình (LMIC) ở Đông Á và '
  'Thái Bình Dương. Tổng cộng 69 nghiên cứu từ 12 quốc gia: 32 RCT, 31 nghiên cứu trước–sau, '
  '6 đánh giá sau can thiệp. Đăng trên Asia Pacific Journal of Public Health vol. 37(4), '
  'tr. 332–346 (Menon et al. 2025).', align='justify')

P('Phát hiện then chốt (theo abstract Crossref):', bold=True)
P('• Có sự TẬP TRUNG MẤT CÂN ĐỐI vào "individual capacity" (năng lực cá nhân) và "clinical '
  'management" (quản lý lâm sàng).', align='justify')
P('• KHOẢNG TRỐNG ở: (a) thúc đẩy SK dựa CỘNG ĐỒNG; (b) phòng ngừa cấp GIA ĐÌNH; (c) dịch '
  'vụ ĐÁP ỨNG dài hạn.', align='justify')
P('• Phần lớn nghiên cứu tập trung tại Trung Quốc và Đông Nam Á; các nước nhỏ hơn và Thái '
  'Bình Dương có đại diện tối thiểu.', align='justify')

P('Phản biện: Đây là CẢNH BÁO MẠNH cho dự án Việt Nam — không nên chỉ dừng lại ở can thiệp '
  'cấp cá nhân (CBT cho học sinh). Cần thiết kế can thiệp ĐA CẤP: cá nhân + GIA ĐÌNH + '
  'CỘNG ĐỒNG + dài hạn. Phù hợp đặc biệt với phát hiện của Dong, Wang & Lin (2025, PLOS ONE '
  '20(9):e0328785, Bảng 3 tr. 7) — kênh giao tiếp gia đình giảm 78 % nguy cơ trầm cảm '
  '(OR = 0,22) và Chen et al. (2025, J Affect Disord 374:408–432) — emotional functioning '
  'và family support là buffer chính trong COVID era. Gợi ý mạnh: thêm module GIA ĐÌNH vào '
  'thiết kế can thiệp Việt Nam.', italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('2.8. Li et al. (2025) — NMA Bayesian 30 RCT cho lo âu trẻ em', level=2)
P('Phân tích tổng hợp mạng Bayesian 30 RCT về can thiệp cho rối loạn lo âu ở trẻ em. Đăng '
  'trên BMC Psychiatry vol. 25, article 809, tr. 1–14 (Li et al. 2025). PROSPERO '
  'CRD42024587910. Tổng cộng 1.711 người tham gia từ 12 quốc gia.', align='justify')

P('Kết quả SUCRA (Li et al. 2025, Bảng 3 tr. 6):', bold=True)
table(
    ['Liệu pháp', 'Số RCT', 'MD (KTC 95 % CrI)', 'SUCRA', 'Ghi chú'],
    [
        ['ACT', '6', '−3,83 (−9,33 đến 1,51)', '0,69', 'Hạng 1 nhưng CrI bao gồm 0'],
        ['CBT', '16', '−3,64 (−7,36 đến −0,48)', '0,66', 'Hạng 2 — bằng chứng MẠNH NHẤT, CrI không qua 0'],
        ['VRET', '6', '−2,53 (−8,23 đến 3,32)', '0,51', 'Hạng 3, CrI rộng'],
        ['PE', '2', '−2,16 (−9,99 đến 5,52)', '0,51', 'Hạng 4 — chỉ 2 RCT'],
        ['Đối chứng', '—', 'Ref', '—', '—'],
    ],
    widths=[2.0, 1.5, 4.5, 1.5, 6.0]
)
P('Bảng 5. SUCRA của các can thiệp cho lo âu trẻ em. '
  'Nguồn: Li et al. 2025, BMC Psychiatry 25:809, Bảng 3 tr. 6. '
  'ACT = Acceptance and Commitment Therapy; VRET = Virtual Reality Exposure Therapy; '
  'PE = Physical Exercise.', italic=True, size=10)

img('chart7_BMC_NMA_Li.png', width_cm=15.5)

P('Phản biện: Đáng chú ý ACT (Acceptance and Commitment Therapy) có SUCRA cao nhất 0,69 — '
  'BẤT NGỜ vì CBT thường được coi là hàng đầu. Tuy nhiên CrI của ACT bao gồm 0 (−9,33 đến '
  '+1,51), nên kết quả không có ý nghĩa thống kê chắc chắn. CBT xếp hạng 2 nhưng có CrI '
  'KHÔNG bao gồm 0 (−7,36 đến −0,48) và có 16 RCT (so với 6 cho ACT) — bằng chứng vững chắc '
  'nhất. PE chỉ có 2 RCT — cần thêm nghiên cứu. Cho Việt Nam: ưu tiên CBT là cốt lõi, có thể '
  'thử nghiệm ACT bổ trợ (mô hình mới), tích hợp PE vì dễ triển khai tại trường.',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

# ============================================================
# PHẦN III — ÂU - ÚC - MỸ
# ============================================================
doc.add_page_break()
H('PHẦN III — CHÂU ÂU – CHÂU ÚC – CHÂU MỸ', level=1)

P('Phần này tổng hợp 7 công trình từ Anh, Mỹ và phân tích tổng hợp quốc tế đăng trên các '
  'tạp chí phương Tây — bao gồm cả các MA quốc tế dựa chủ yếu trên RCT phương Tây. Đây là '
  'vùng có bằng chứng lâu dài và phong phú nhất, cung cấp nền tảng phương pháp luận và các '
  'mô hình "gold standard" cho thiết kế can thiệp tại Việt Nam.', align='justify', italic=True)

H('3.1. Tổng quan 7 nghiên cứu can thiệp Châu Âu – Úc – Mỹ', level=2)
table(
    ['#', 'Tác giả / Năm', 'Vùng', 'Loại nghiên cứu', 'Mẫu / Cỡ mẫu', 'Phát hiện then chốt'],
    [
        ['44', 'Walder et al. 2025', 'Quốc tế (JMIR)', 'MA 21 RCT', 'Trẻ + VTN + thanh niên SAD', 'DMHI g = 0,508; SAD-specific 0,878'],
        ['49', 'Bress et al. 2024', 'Mỹ', 'RCT', '59 thanh niên 18–25 tuổi', 'Maya app: HAM-A d = 1,04 (12 tuần)'],
        ['48', 'Brown & Carter 2025', 'Anh quốc', 'Editorial Q1', '—', 'Mindfulness phổ quát 8.376 HS thất bại; BESST 900 HS dương tính'],
        ['50', 'Cao et al. 2025', 'Quốc tế', 'SR + MA RCT', 'Trẻ + VTN', 'Resilience trường: hiệu ứng nhỏ–TB; heterogeneity cao'],
        ['56*', 'Zhang, Liang & Kang 2026', 'Quốc tế', 'Bayesian MA 31 RCT', '19.865 HS nguy cơ thấp', 'CBT phổ quát: hiệu ứng NHỎ; chất lượng nền THẤP'],
        ['28', 'Zugman et al. 2024', 'Mỹ', 'Review treatment', 'VTN', 'Tổng quan CBT + SSRI; trích dẫn CAMS NEJM'],
        ['—', 'Walkup et al. 2008', 'Mỹ', 'RCT đa trung tâm', '488 trẻ 7–17 tuổi', 'CBT + SSRI 80,7 %; CBT đơn 59,7 %; SSRI 54,9 %; placebo 23,7 %'],
    ],
    widths=[0.8, 3.2, 1.8, 2.4, 3.1, 4.2]
)
P('Bảng 6. Tổng quan 7 nghiên cứu can thiệp Châu Âu – Úc – Mỹ. '
  '* Bài 56 chỉ có abstract (paywall Wiley).', italic=True, size=10)

H('3.2. Walder et al. (2025) — Meta-analysis 21 RCT về DMHI cho lo âu xã hội', level=2)
P('Phân tích tổng hợp 21 RCT về Can thiệp Sức khoẻ Tâm thần Số (Digital Mental Health '
  'Interventions — DMHI) cho phòng ngừa và điều trị rối loạn lo âu xã hội ở trẻ em, vị thành '
  'niên và thanh niên. Đăng trên Journal of Medical Internet Research (JMIR — Q1, IF ≈ 7,4) '
  '(Walder et al. 2025). PROSPERO CRD42023424181.', align='justify')

P('Kết quả Hedges g cho các phân nhóm (Walder et al. 2025, Bảng 2 tr. 18):', bold=True)
table(
    ['Phân nhóm', 'Hedges g', 'KTC 95 %', 'Ý nghĩa lâm sàng'],
    [
        ['DMHI vs bất kỳ đối chứng', '0,508', '0,308 – 0,707', 'TRUNG BÌNH — đáng kể'],
        ['DMHI vs danh sách chờ (WL)', '0,576', '0,343 – 0,809', 'Mạnh hơn vs WL'],
        ['DMHI dựa trên CBT', '0,610', '0,361 – 0,859', 'Nền tảng CBT TỐT NHẤT'],
        ['DMHI có HƯỚNG DẪN người', '0,825', '0,425 – 1,224', 'LỚN — hướng dẫn quan trọng'],
        ['DMHI thiết kế RIÊNG cho SAD', '0,878', '0,469 – 1,278', 'LỚN NHẤT — SAD-specific'],
        ['DMHI KHÔNG hướng dẫn', '~ 0,3 – 0,4', '—', 'Nhỏ — kém hiệu quả'],
        ['Sau xét thiên lệch xuất bản', '0,506', '0,308 – 0,707', 'Vẫn đáng kể — kết quả robust'],
    ],
    widths=[6.0, 2.0, 3.0, 4.5]
)
P('Bảng 7. Hiệu lực DMHI cho lo âu xã hội theo phân nhóm. '
  'Nguồn: Walder et al. 2025, JMIR, Bảng 2, tr. 18.', italic=True, size=10)

img('chart4_DMHI_Walder.png', width_cm=15.5)

P('Phản biện: Đây là bằng chứng MẠNH NHẤT về DMHI / iCBT cho SAD ở trẻ và VTN trong toàn bộ '
  'tài liệu. Hai phát hiện cốt lõi: (1) thiết kế ĐẶC THÙ cho SAD (g = 0,878) hiệu lực hơn '
  'gần gấp đôi DMHI tổng quát (g = 0,508); (2) có HƯỚNG DẪN người (g = 0,825) hiệu lực gần '
  'gấp đôi không có hướng dẫn (g ≈ 0,3–0,4). Phù hợp với Xian, Zhang & Jiang 2024 (NMA hạng '
  '1 cho iCBT) và Sasaki et al. 2024 (RCT Nhật iCBT dương tính). Cho Việt Nam: nếu phát '
  'triển app tiếng Việt, BẮT BUỘC phải (a) dành riêng cho SAD / lo âu, (b) có thành phần '
  'hỗ trợ người (chat, video call định kỳ, hoặc nhân viên hướng dẫn). Hai điều kiện này '
  'tăng hiệu lực gấp ~ 3 lần.', italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('3.3. Bress et al. (2024) — RCT ứng dụng Maya tại JAMA Network Open', level=2)
P('Thử nghiệm lâm sàng ngẫu nhiên đăng tại JAMA Network Open vol. 7(8), e2428372, tr. 1–15 '
  '(Bress et al. 2024). DOI 10.1001/jamanetworkopen.2024.28372. ClinicalTrials.gov '
  'NCT05130281. Đối tượng: 59 thanh niên 18–25 tuổi có rối loạn lo âu (GAD 56 %, SAD 41 %), '
  'tuyển qua quảng cáo trực tuyến.', align='justify')

P('Can thiệp: ứng dụng Maya — CBT tự hướng dẫn 12 phiên / 6 tuần. 3 nhóm khuyến khích dựa '
  'trên tin nhắn (gain-frame, loss-frame, social support). Theo dõi đến tuần 12. Đo bằng '
  'HAM-A (chính), ASI, LSAS, HDRS-24, uMARS.', align='justify')

P('Kết quả thay đổi HAM-A từ ban đầu (Bress et al. 2024, Bảng 2 tr. 7):', bold=True)
table(
    ['Thời điểm', 'Chênh lệch TB', 'KTC 95 %', "Cohen's d", 'p'],
    [
        ['Tuần 3 (giữa)', '−3,20', '−4,76 đến −1,64', '0,64 (TRUNG BÌNH)', '< 0,001'],
        ['Tuần 6 (kết thúc)', '−5,64', '−7,23 đến −4,05', '0,94 (LỚN)', '< 0,001'],
        ['Tuần 12 (theo dõi)', '−5,67', '−7,29 đến −4,04', '1,04 (RẤT LỚN)', '< 0,001'],
    ],
    widths=[3.5, 3.0, 3.0, 3.5, 2.5]
)
P('Bảng 8. Hiệu quả ứng dụng Maya trên thang HAM-A. '
  'Nguồn: Bress et al. 2024, JAMA Network Open 7(8):e2428372, Bảng 2, tr. 7.', italic=True, size=10)

img('chart6_Maya_App_Bress.png', width_cm=15.5)

P('Đặc biệt đáng chú ý — hiệu lực TĂNG theo thời gian: Cohen d tăng từ 0,64 (tuần 3) lên '
  '0,94 (tuần 6) và 1,04 (tuần 12). Hiệu quả không chỉ bền mà còn tiếp tục cải thiện sau khi '
  'dừng can thiệp. Trên thang ASI (anxiety sensitivity), Cohen d = 0,93 ở tuần 6; trên LSAS '
  '(lo âu xã hội), d = 1,07 ở tuần 12 (Bảng 3 tr. 8 cùng tài liệu).', align='justify')

P('Phản biện: RCT đăng trên JAMA Network Open (Q1, IF ≈ 13,8) — bằng chứng chất lượng cực '
  'cao. Hiệu lực Cohen d = 1,04 là RẤT LỚN cho can thiệp app — gần như tương đương CBT mặt '
  'đối mặt. Hai lưu ý: (1) đối tượng 18–25 tuổi (THANH NIÊN, không phải VTN học sinh) — cần '
  'kiểm chứng ở tuổi nhỏ hơn; (2) quảng cáo trực tuyến tuyển mộ → mẫu có thể tự chọn cao, '
  'có động lực sẵn → kết quả có thể không generalizable cho dân số chung. Cho Việt Nam: '
  'mô hình app CBT có hỗ trợ tin nhắn rất khả thi cho VTN cấp 3 + SV, cần RCT thử nghiệm.',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('3.4. Brown & Carter (2025) — Editorial UK về can thiệp tại trường', level=2)
P('Editorial đăng trên Journal of Mental Health vol. 34(4), tr. 357–361 (Brown & Carter '
  '2025). Tác giả chính June S.L. Brown — chuyên gia hàng đầu UK về can thiệp tâm lý cộng '
  'đồng. Bài tổng hợp các mô hình can thiệp tại trường ở Anh và đề xuất hướng đi.', align='justify')

table(
    ['Mô hình', 'Người cung cấp', 'Hiệu quả', 'Bằng chứng nguồn'],
    [
        ['Mindfulness phổ quát', 'Giáo viên', 'THẤT BẠI', '8.376 HS / 85 trường (Kuyken 2022)'],
        ['BESST (CBT self-referral)', 'Chuyên gia tâm lý', 'DƯƠNG TÍNH', '900 HS / 57 trường (Brown 2024)'],
        ['MHST (Mental Health Support Team)', 'Thạc sĩ trị liệu', 'HỨA HẸN', 'Mới triển khai NHS UK'],
        ['PLACES (self-referral)', 'Đa cấp', 'HỨA HẸN', 'Dùng từ thường ngày, giảm kỳ thị'],
    ],
    widths=[4.5, 3.5, 3.0, 4.5]
)
P('Bảng 9. Bốn mô hình can thiệp tại trường ở Anh. '
  'Nguồn: Brown & Carter 2025, J Mental Health 34(4), tr. 358–360.', italic=True, size=10)

P('Năm bài học từ Editorial này (Brown & Carter 2025, tr. 358–361):', bold=True)
P('1. UNIVERSAL MINDFULNESS thất bại do engagement thấp ở học sinh (Kuyken 2022, n = 8.376).', align='justify')
P('2. TARGETED CBT (học sinh có triệu chứng tự đăng ký) thành công (BESST trial 900 HS).', align='justify')
P('3. Chuyên gia lâm sàng > giáo viên về hiệu quả, nhưng cả hai đều khả thi với đào tạo phù hợp.', align='justify')
P('4. CO-DESIGN (cho học sinh tham gia thiết kế can thiệp) tăng engagement đáng kể — đặc biệt '
  'với học sinh từ các nhóm thiểu số.', align='justify')
P('5. Ngôn ngữ THƯỜNG NGÀY ("căng thẳng" thay vì "trầm cảm") giảm kỳ thị và tăng tiếp cận.', align='justify')

P('Phản biện: Editorial — không phải SR/MA gốc, là ý kiến chuyên gia + tổng hợp bằng chứng. '
  'Tuy nhiên cung cấp khung phân loại RẤT HỮU ÍCH. Cảnh báo về universal interventions là '
  'QUAN TRỌNG cho thiết kế đề cương Việt Nam. Phù hợp với Zhang, Liang & Kang 2026 '
  '(Bayesian MA — universal CBT hiệu ứng nhỏ). Cho Việt Nam: nên áp dụng mô hình BESST '
  '(targeted + self-referral) thay vì universal, kết hợp PLACES language để giảm kỳ thị.',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('3.5. Cao et al. (2025) — SR + MA về can thiệp resilience tại trường', level=2)
P('Tổng quan hệ thống và phân tích tổng hợp các RCT về can thiệp khả năng phục hồi '
  '(resilience) tại trường cho trẻ em + vị thành niên. Đăng trên Frontiers in Psychiatry '
  'vol. 16, article 1594658, tr. 1–15 (Cao et al. 2025). Q1, IF ≈ 4,7. Open Access.', align='justify')

P('Resilience được định nghĩa là khả năng thích ứng tốt trước nghịch cảnh, căng thẳng hoặc '
  'sang chấn. Khác với CBT (giảm triệu chứng), can thiệp resilience nhắm tăng cường yếu tố '
  'BẢO VỆ: tự trọng, kỹ năng ứng phó, kết nối xã hội, lạc quan. Đây là tiếp cận từ tâm lý '
  'tích cực (positive psychology).', align='justify')

P('Phương pháp: Tìm kiếm hệ thống các RCT, đánh giá chất lượng bằng Cochrane Risk of Bias, '
  'random-effects meta-analysis. Kết quả: can thiệp resilience tại trường CÓ HIỆU QUẢ tăng '
  'resilience và giảm triệu chứng SKTT. Tuy nhiên kích thước hiệu ứng NHỎ – TRUNG BÌNH và '
  'heterogeneity CAO (Cao et al. 2025, tr. 6–9).', align='justify')

P('Phản biện: SR + MA RCTs trên Frontiers Q1 — bằng chứng tốt. Effect size nhỏ – TB cho '
  'thấy resilience không thay thế được CBT, nhưng là YẾU TỐ BẢO VỆ bổ trợ. Phù hợp với phát '
  'hiện của Trần Thảo Vi (2025, Tâm Lý Học) — lạc quan trung gian quan hệ lo âu – trầm cảm '
  '(β gián tiếp = −0,24). Phù hợp với Ireland MyWorld Survey (Fitzgerald et al. 2024) — '
  'resilience + tự trọng quan trọng tăng theo thời gian (2012 → 2019). Cho Việt Nam: tích '
  'hợp module resilience (3 buổi) vào can thiệp CBT chính.', italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('3.6. Zhang, Liang & Kang (2026) — Bayesian MA 31 RCT về CBT trường dài hạn', level=2)
P('Bayesian Meta-Analysis 31 RCT về CBT phổ quát tại trường cho trầm cảm và lo âu ở trẻ em '
  'và VTN nguy cơ thấp (low-risk). Tổng N = 19.865 — mẫu lớn nhất hiện nay cho can thiệp '
  'CBT trường. Đăng trên Journal of Clinical Psychology vol. 82, tr. 248–259, online '
  '11/2025, in print 03/2026 (Zhang, Liang & Kang 2026). Bài này hiện chỉ có abstract qua '
  'paywall Wiley.', align='justify')

P('Kết quả chính (theo abstract Crossref):', bold=True)
P('• Cải thiện CÓ Ý NGHĨA THỐNG KÊ nhưng KHIÊM TỐN trên triệu chứng trầm cảm.', align='justify')
P('• Giảm NHỎ trên triệu chứng lo âu.', align='justify')
P('• Hiệu quả DUY TRÌ tới 1 năm sau can thiệp.', align='justify')
P('• Kết luận của tác giả: "Chất lượng bằng chứng nền RẤT THẤP khiến phát hiện này CHƯA ĐỦ '
  'vững chắc để hỗ trợ triển khai rộng rãi tại thời điểm này."', align='justify')

P('Phản biện: CẢNH BÁO ĐẶC BIỆT QUAN TRỌNG — đây là MA có cỡ mẫu lớn nhất (n = 19.865) '
  'nhưng kết luận thận trọng. Phù hợp với Brown & Carter 2025 (mindfulness universal thất '
  'bại UK). Hai bằng chứng độc lập đều cảnh báo: KHÔNG nên triển khai universal CBT phổ '
  'quát tại trường mà chưa có bằng chứng chất lượng cao hơn. Cho Việt Nam: nên thiết kế '
  'can thiệp TARGETED (cho học sinh có triệu chứng lo âu / trầm cảm cụ thể) thay vì cho '
  'toàn bộ học sinh — tăng dilution + giảm engagement.',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

H('3.7. Zugman et al. (2024) và CAMS (Walkup et al. 2008) — Bằng chứng vàng về CBT + SSRI', level=2)
P('Bài tổng quan đăng trên American Journal of Psychiatry vol. 181(3), tr. 189–200, March '
  '2024 (Zugman et al. 2024). DOI 10.1176/appi.ajp.20231037. AJP — Q1, IF ≈ 18, tạp chí '
  'tâm thần hàng đầu thế giới. Tác giả từ National Institute of Mental Health (NIMH/NIH) — '
  'cơ quan nghiên cứu SKTT hàng đầu Mỹ.', align='justify')

P('Bài tập trung vào 3 rối loạn lo âu trẻ em chính: GAD (rối loạn lo âu tổng quát), '
  'separation anxiety (lo âu chia ly), và social anxiety disorder (lo âu xã hội). Tổng quan '
  'tài liệu CBT, SSRIs, kết hợp + hướng tương lai (ABMT, DCS, VR, biomarker).', align='justify')

P('Trích dẫn cốt lõi từ thử nghiệm CAMS (Walkup et al. 2008, NEJM 359:2753–2766) — được '
  'Zugman et al. 2024 trình bày ở tr. 191:', bold=True)

table(
    ['Nhánh điều trị', 'Tỷ lệ đáp ứng', 'NNT', 'Ý nghĩa'],
    [
        ['CBT + SSRI (Sertraline) — KẾT HỢP', '80,7 %', '~ 2', 'TỐT NHẤT — kết hợp tâm lý + thuốc'],
        ['CBT đơn thuần', '59,7 %', '~ 3', 'TỐT — phương pháp tâm lý đơn lẻ'],
        ['SSRI đơn thuần (Sertraline)', '54,9 %', '~ 4', 'KHÁ — thuốc đơn lẻ'],
        ['Placebo', '23,7 %', '—', 'BASELINE'],
    ],
    widths=[5.5, 2.5, 1.5, 6.0]
)
P('Bảng 10. Đáp ứng điều trị nghiên cứu CAMS (n = 488 trẻ 7–17 tuổi, RCT đa trung tâm). '
  'Nguồn: Walkup et al. 2008, NEJM 359:2753–66 — trích dẫn trong Zugman et al. 2024, '
  'Am J Psychiatry 181(3), tr. 191. NNT = Number Needed to Treat.', italic=True, size=10)

img('chart5_CAMS_Walkup.png', width_cm=15.5)

P('Phản biện: CAMS là RCT đa trung tâm 488 trẻ 7–17 tuổi — tiêu chuẩn vàng (gold standard) '
  'cho điều trị RLLA nhi khoa. Mặc dù từ 2008, CAMS vẫn được trích dẫn rộng rãi và là cơ sở '
  'cho hướng dẫn lâm sàng quốc tế. Phát hiện cốt lõi: kết hợp CBT + SSRI hiệu quả nhất, '
  'nhưng CBT đơn thuần (không thuốc) cũng đạt 59,7 % — gấp 2,5 lần placebo. Hạn chế: chỉ Mỹ; '
  'có thể không generalizable trực tiếp cho LMIC do chi phí thuốc + chuyên gia + cân nhắc '
  'đạo đức về SSRI cho VTN. Cho Việt Nam: nên ƯU TIÊN CBT đơn thuần (không SSRI) — phù hợp '
  'với năng lực chuyên môn, chi phí và đạo đức kê đơn cho VTN tại VN.',
  italic=True, color=RGBColor(0xCC,0,0), align='justify')

# ============================================================
# PHẦN IV — TỔNG KẾT & KHUYẾN NGHỊ
# ============================================================
doc.add_page_break()
H('PHẦN IV — TỔNG KẾT VÀ KHUYẾN NGHỊ CHO VIỆT NAM', level=1)

H('4.1. So sánh ba vùng địa lý', level=2)
table(
    ['Tiêu chí', 'Việt Nam', 'Châu Á (ngoài VN)', 'Châu Âu – Úc – Mỹ'],
    [
        ['Số bài can thiệp', '1', '7', '7'],
        ['Đối tượng VTN', 'KHÔNG (người lớn)', 'CÓ (đa số)', 'CÓ (đa số)'],
        ['Loại NC chủ yếu', 'Trước–sau, không đối chứng', 'RCT + NMA + cluster RCT', 'RCT + MA + Bayesian MA'],
        ['Phương pháp chính', 'Thư giãn–Luyện tập (Yoga + thở)', 'CBT + iCBT + Mobile + CA-CBT', 'CBT + iCBT + DMHI + Mindfulness'],
        ['Bối cảnh thực hiện', 'Bệnh viện (nội trú)', 'Trường + BV + Online', 'Trường + Online + Phòng khám'],
        ['Bằng chứng cao nhất', 'Yếu (1 luận án)', 'Mạnh (Xian NMA)', 'Rất mạnh (CAMS NEJM, Bayesian MA)'],
        ['Khoảng trống', 'CỰC LỚN — 0 RCT cho VTN', 'Vừa — thiếu cộng đồng/gia đình', 'Nhỏ — đã có nhiều bằng chứng'],
    ],
    widths=[3.5, 4.0, 4.0, 4.0]
)
P('Bảng 11. So sánh tổng hợp ba vùng địa lý.', italic=True, size=10)

H('4.2. Xếp hạng 13 phương pháp can thiệp dựa trên tổng hợp bằng chứng', level=2)
table(
    ['Hạng', 'Phương pháp', 'Bằng chứng tốt nhất', 'Phù hợp VN', 'Khuyến nghị VN'],
    [
        ['1', 'CBT + SSRI kết hợp', 'CAMS NEJM 80,7 %', 'Hạn chế (chi phí, đạo đức)', 'Chỉ ca nặng có chỉ định'],
        ['2', 'CBT cá nhân / nhóm', 'CAMS 59,7 %; Li 2025 SUCRA 0,66', 'CỐT LÕI', 'CHÍNH — đầu tư đào tạo'],
        ['3', 'iCBT cho SAD', 'Xian 2024 NMA hạng 1; Walder g = 0,878', 'CAO — VTN số hoá', 'Phát triển app tiếng Việt'],
        ['4', 'gCBT (CBT nhóm)', 'Xian 2024 NMA hạng 2 chức năng', 'CAO — chi phí thấp', 'Triển khai trường'],
        ['5', 'CBT do GV (school-based)', 'De Silva 2024 cluster RCT 720 HS', 'CAO — LMIC khả thi', 'Đào tạo GV'],
        ['6', 'CA-CBT (thích ứng văn hoá)', 'Praptomojati 2024 SR ĐNA', 'CỰC CAO', 'Phát triển bản VN'],
        ['7', 'PE (hoạt động thể chất)', 'Li 2025 SUCRA 0,51', 'CAO — sẵn có', 'Tích hợp giờ thể dục'],
        ['8', 'Mobile App CBT (Maya-style)', 'Bress 2024 d = 1,04', 'CAO', 'RCT pilot'],
        ['9', 'Resilience training', 'Cao 2025 MA Frontiers', 'CAO', 'Module bổ trợ'],
        ['10', 'Thư giãn – Luyện tập (Yoga)', 'Trần Nguyễn Ngọc 2018', 'CAO — đã thử ở VN', 'Module bổ trợ'],
        ['11', 'Mobile CBT (trầm cảm)', 'Qiaochu 2025 87,5 %', 'CAO cho trầm cảm', 'Mạnh trầm cảm hơn lo âu'],
        ['12', 'PLACES self-referral', 'Brown & Carter 2025 BESST', 'TB — cần adapt', 'Pilot thử nghiệm'],
        ['13', 'Mindfulness phổ quát', 'Brown 2025: thất bại 8.376 HS', 'KHÔNG nên', 'TRÁNH triển khai universal'],
    ],
    widths=[1.0, 4.0, 4.0, 3.5, 3.0]
)
P('Bảng 12. Xếp hạng 13 phương pháp can thiệp tâm lý cho RLLA ở VTN. '
  'Tổng hợp từ 15 bài nghiên cứu trong báo cáo này.', italic=True, size=10)

H('4.3. Năm khoảng trống nghiên cứu can thiệp tại Việt Nam', level=2)
P('Sau khi tổng hợp 15 bài, có thể xác định 5 khoảng trống lớn nhất cho lĩnh vực can thiệp '
  'RLLA ở VTN tại Việt Nam:', bold=True, align='justify')
P('1. THIẾU TUYỆT ĐỐI: 0 RCT can thiệp tâm lý cho VTN có RLLA tại Việt Nam trong hệ thống '
  '58 bài. Trần Nguyễn Ngọc 2018 là người lớn nội trú.', align='justify')
P('2. Phát triển app iCBT / Mobile CBT TIẾNG VIỆT — phù hợp xếp hạng số 3, 4, 8 (Xian '
  '2024, Walder 2025, Bress 2024).', align='justify')
P('3. Đào tạo GIÁO VIÊN cung cấp CBT tại trường — mô hình De Silva 2024 (Sri Lanka) khả thi '
  'cho LMIC.', align='justify')
P('4. CA-CBT (CBT thích ứng văn hoá Á Đông) phiên bản VN — yếu tố Phật giáo, gia đình mở '
  'rộng, "fear of letting down" (Dong, Wang & Lin 2025).', align='justify')
P('5. Can thiệp ĐA CẤP: cá nhân + GIA ĐÌNH + cộng đồng — lấp khoảng trống Menon 2025 và '
  'phát hiện kênh giao tiếp gia đình của Dong, Wang & Lin 2025 (OR = 0,22; tr. 7).', align='justify')

H('4.4. Đề xuất thiết kế can thiệp cho đề cương Việt Nam (giai đoạn 2)', level=2)
P('Dựa trên tổng hợp 15 bài và 5 khoảng trống trên, đề xuất can thiệp PHỐI HỢP 12 tuần '
  'TARGETED cho học sinh THCS / THPT có triệu chứng lo âu (sàng lọc bằng GAD-7 ≥ 8 '
  'hoặc DASS-Y lo âu ≥ 8):', bold=True, align='justify')

table(
    ['Thành phần', 'Số buổi', 'Cơ sở bằng chứng', 'Áp dụng VN'],
    [
        ['1. CBT NHÓM TARGETED', '8 buổi × 60 phút', 'CAMS Walkup 2008 NEJM; Li 2025 SUCRA 0,66; De Silva 2024 RCT', 'Cốt lõi can thiệp; do GV/cố vấn tâm lý đã đào tạo cung cấp'],
        ['2. Module GIAO TIẾP GIA ĐÌNH', '4 buổi (cha mẹ + HS)', 'Dong, Wang & Lin 2025 OR = 0,22 (mới); Menon 2025 gap gia đình', 'Tập huấn kỹ năng tâm sự; lắng nghe; không phán xét'],
        ['3. Module RESILIENCE', '3 buổi', 'Cao 2025 Frontiers MA; Trần Thảo Vi 2025 (lạc quan)', 'Lạc quan, kết nối XH, tự nhận thức'],
        ['4. PE + Thư giãn – Thở', '8 buổi (giờ thể dục)', 'Li 2025 PE SUCRA 0,51; Trần Nguyễn Ngọc 2018', 'Tích hợp giờ thể dục; Yoga đã thử nghiệm ở VN'],
        ['5. App iCBT TIẾNG VIỆT', 'Suốt 12 tuần', 'Walder 2025 g = 0,878 (SAD-specific + guided); Sasaki 2024', 'Bắt buộc có hỗ trợ người (chat); thiết kế cho lo âu'],
    ],
    widths=[3.5, 2.5, 5.0, 4.5]
)
P('Bảng 13. Đề xuất can thiệp 5 thành phần cho đề cương VN giai đoạn 2.', italic=True, size=10)

P('Cơ chế tiếp cận: PLACES self-referral (Brown & Carter 2025) — học sinh tự đăng ký dùng '
  'từ thường ngày ("căng thẳng", "lo lắng"), kết hợp sàng lọc chủ động để tăng tiếp cận.', bold=True, align='justify')

H('4.5. Kết luận', level=2)
P('Việt Nam đang có khoảng trống RẤT LỚN về nghiên cứu can thiệp tâm lý cho VTN có RLLA — '
  '0 RCT trong hệ thống 58 bài. Trần Nguyễn Ngọc 2018 là cơ sở duy nhất từ Việt Nam nhưng '
  'đối tượng người lớn nội trú. Bằng chứng quốc tế (đặc biệt từ Châu Á) cho thấy CBT (cá '
  'nhân, nhóm, internet) là phương pháp hiệu quả nhất; mô hình Sri Lanka (GV cung cấp) khả '
  'thi cho LMIC; iCBT phù hợp VTN số hoá; PHẢI tránh universal interventions phổ quát; cần '
  'tích hợp can thiệp GIA ĐÌNH (phát hiện mới Dong, Wang & Lin 2025). Đề cương Việt Nam '
  'nên thiết kế can thiệp TARGETED 12 tuần đa thành phần — đây sẽ là RCT đầu tiên loại này '
  'tại Việt Nam.', bold=True, italic=True, align='justify')

# ============================================================
# TÀI LIỆU THAM KHẢO
# ============================================================
doc.add_page_break()
H('TÀI LIỆU THAM KHẢO', level=1)

P('Bài 28 — Zugman, A., Jett, J. E., Antonacci, C., Winkler, A. M., & Pine, D. S. (2024). '
  'Current and future approaches to pediatric anxiety disorder treatment. American Journal '
  'of Psychiatry, 181(3), 189–200. DOI 10.1176/appi.ajp.20231037', align='justify')

P('Bài 29 — Li, L., Li, Q., Wang, J., Fu, Q., Cui, M., Wang, Y., et al. (2025). Effects of '
  'different interventions on anxiety disorders in children and adolescents: A systematic '
  'review and Bayesian network meta-analysis. BMC Psychiatry, 25, 809. '
  'DOI 10.1186/s12888-025-07227-y. PROSPERO CRD42024587910.', align='justify')

P('Bài 41 — Praptomojati, A., & Hartanto, A. (2024). A systematic review of culturally '
  'adapted cognitive behavioral therapy (CA-CBT) for anxiety disorders in Southeast Asia. '
  'Asian Journal of Psychiatry, 92, 103896. DOI 10.1016/j.ajp.2023.103896', align='justify')

P('Bài 42 — De Silva, S., Bowers, A., Wijesinghe, M., Gamage, A., & Hapangama, A. (2024). '
  'Effectiveness of a cognitive behavioural therapy (CBT)-based intervention for reducing '
  'anxiety among adolescents in the Colombo District, Sri Lanka: A cluster randomised '
  'controlled trial. Child and Adolescent Psychiatry and Mental Health, 18, 108. '
  'DOI 10.1186/s13034-024-00799-9', align='justify')

P('Bài 43 — Xian, J., Zhang, Y., & Jiang, B. (2024). Psychological interventions for social '
  'anxiety disorder in children and adolescents: A systematic review and network meta-'
  'analysis. Journal of Affective Disorders, 365, 614–627. DOI 10.1016/j.jad.2024.08.097. '
  'PROSPERO CRD42023476829.', align='justify')

P('Bài 44 — Walder, N., Frey, A., Berger, T., et al. (2025). Digital mental health '
  'interventions for prevention and treatment of social anxiety disorder in children, '
  'adolescents and young adults: Systematic review and meta-analysis. Journal of Medical '
  'Internet Research, e67067. DOI 10.2196/preprints.67067. PROSPERO CRD42023424181.', align='justify')

P('Bài 48 — Brown, J. S. L., & Carter, B. (2025). School-based interventions for depression '
  'and anxiety in UK [Editorial]. Journal of Mental Health, 34(4), 357–361. '
  'DOI 10.1080/09638237.2025.2512332', align='justify')

P('Bài 49 — Bress, J. N., Falk, A., Schier, M. M., et al. (2024). Efficacy of a mobile '
  'app-based intervention for young adults with anxiety disorders: A randomized clinical '
  'trial. JAMA Network Open, 7(8), e2428372. DOI 10.1001/jamanetworkopen.2024.28372. '
  'ClinicalTrials.gov NCT05130281.', align='justify')

P('Bài 50 — Cao, C., et al. (2025). School-based interventions for resilience in children '
  'and adolescents: A systematic review and meta-analysis of randomized controlled trials. '
  'Frontiers in Psychiatry, 16, 1594658. DOI 10.3389/fpsyt.2025.1594658', align='justify')

P('Bài 51 — Sasaki, N., et al. (2024). Effectiveness of unguided internet-based cognitive '
  'behavioral therapy for subthreshold social anxiety disorder in adolescents and young '
  'adults: Multicenter randomized controlled trial. JMIR Pediatrics and Parenting, 7, '
  'e55786. DOI 10.2196/55786. UMIN000049768.', align='justify')

P('Bài 56* — Zhang, X., Liang, Z., & Kang, J. (2026). Long-term effects of school-based '
  'CBT in low-risk children and adolescents: A Bayesian meta-analysis. Journal of Clinical '
  'Psychology, 82, 248–259. DOI 10.1002/jclp.70069 [abstract-only].', align='justify')

P('Bài 57* — Qiaochu, Z., & Wang, Y. (2025). Effectiveness of mobile-based cognitive '
  'behavioural therapy for depression and anxiety in children and young people: A '
  'systematic review of randomized controlled trials. Clinical Psychology and Psychotherapy, '
  '32(6), e70173. DOI 10.1002/cpp.70173 [abstract-only].', align='justify')

P('Bài 58* — Menon, V., Coppard, M., McEwen, S., Romero, L., Kennedy, E., & Azzopardi, P. '
  '(2025). Evaluated interventions targeting the mental health and psychosocial wellbeing '
  'of children and adolescents: A scoping review focused on low- and middle-income '
  'countries in East Asia and the Pacific. Asia Pacific Journal of Public Health, 37(4), '
  '332–346. DOI 10.1177/10105395241313154 [abstract-only].', align='justify')

P('Trần Nguyễn Ngọc (2018). Đánh giá hiệu quả điều trị rối loạn lo âu lan toả bằng liệu '
  'pháp thư giãn – luyện tập [Luận án Tiến sĩ Y học, mã số 62720148]. Trường Đại học Y Hà '
  'Nội — Viện Sức khoẻ Tâm thần, Bệnh viện Bạch Mai. Người hướng dẫn: PGS. TS. Nguyễn Kim Việt. '
  '177 trang.', align='justify')

P('Walkup, J. T., Albano, A. M., Piacentini, J., Birmaher, B., Compton, S. N., Sherrill, '
  'J. T., et al. (2008). Cognitive behavioral therapy, sertraline, or a combination in '
  'childhood anxiety. New England Journal of Medicine, 359(26), 2753–2766. '
  'DOI 10.1056/NEJMoa0804633. [Trích dẫn trong Zugman et al. 2024, AJP 181(3), tr. 191].', align='justify')

P('Bài 54 (đối chiếu) — Dong, T., Wang, Y., & Lin, Y. (2025). Prevalence and determinants '
  'of depression, anxiety, and stress among secondary school students. PLOS ONE, 20(9), '
  'e0328785. DOI 10.1371/journal.pone.0328785', align='justify')

P('Bài 55 (đối chiếu) — Chen, H., Wang, Q., Zhu, J., Zhu, Y., Yang, F., Hui, J., Tang, X., '
  '& Zhang, T. (2025). Protective and risk factors of anxiety in children and adolescents '
  'during COVID-19: A systematic review and three level meta-analysis. Journal of '
  'Affective Disorders, 374, 408–432. DOI 10.1016/j.jad.2025.01.029', align='justify')

P('* Các bài đánh dấu * chỉ có thông tin abstract qua paywall (chưa có toàn văn).',
  italic=True, size=10)

P('---', italic=True, align='center')
P('Tổng hợp 15 công trình can thiệp + 2 bài đối chiếu | Phong cách Công Thị Hằng v5\n'
  'Cập nhật: 10/04/2026 | Dự án Lo âu — Nhóm nghiên cứu Tâm lý học',
  italic=True, size=10, align='center')

doc.save(OUT)
print(f'Saved: {OUT}')
d2 = Document(OUT)
print(f'Total paragraphs: {len(d2.paragraphs)}, tables: {len(d2.tables)}')
print(f'Word count estimate: {sum(len(p.text.split()) for p in d2.paragraphs)} words')
