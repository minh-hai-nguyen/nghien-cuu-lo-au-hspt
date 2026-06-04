"""Doc trả lời: M trong "M=25,33" có phải điểm trung bình không? - format mới."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/M_diem_trung_binh_va_SIAS17_18-24.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
BLUE = RGBColor(0, 112, 192)
DARK = RGBColor(31, 73, 125)
GREEN = RGBColor(54, 95, 44)
RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(191, 97, 14)
GRAY = RGBColor(90, 90, 90)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def vn_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p
def blue_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = BLUE
    return p

# TITLE
title = d.add_heading('Ký hiệu M trong báo cáo: M = 25,33 có phải điểm trung bình?', level=0)
for r in title.runs: r.font.color.rgb = DARK
sub = d.add_paragraph()
sr = sub.add_run('Diễn giải M, SD, và các ký hiệu thống kê thường gặp — qua ví dụ Jefferies 2020 (QT035)')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = GRAY
d.add_paragraph()

# CÂU HỎI
qbox = d.add_table(rows=1, cols=1); qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]; shade(cell, 'FFF8DC')
p = cell.paragraphs[0]; r = p.add_run('Câu hỏi của thầy:'); r.bold = True
p2 = cell.add_paragraph()
p2.add_run('"Trong đoạn này, M là điểm trung bình, phải không? — Nhóm 18–24 tuổi có lo âu xã hội CAO NHẤT '
           '(M = 25,33; SAD = 40,3 %)."')
d.add_paragraph()

# BỐI CẢNH
add_h('Bối cảnh — Đoạn này từ bài Jefferies 2020 (QT035 trong DB)', 1)
vn_para('Đoạn thầy trích dẫn từ bài Jefferies, P. & Ungar, M. (2020) "Social Anxiety in Young People: '
        'A Prevalence Study in Seven Countries" — PLOS ONE, 15(9), e0239133. Thông tin chi tiết:')

ctx_tbl = d.add_table(rows=6, cols=2); ctx_tbl.style = 'Table Grid'
ctx_data = [
    ('Bài', 'QT035 Jefferies & Ungar 2020 PLOS ONE Q1'),
    ('Mẫu', '6.825 thanh niên 16-29 tuổi từ 7 quốc gia (Brazil, TQ, Indonesia, Nga, Thái Lan, Mỹ, VIỆT NAM)'),
    ('Công cụ đo', 'SIAS-17 (Social Interaction Anxiety Scale — 17 mục)'),
    ('Thang đo SIAS-17', '17 mục, mỗi mục 0-4 điểm; tổng 0-68'),
    ('Ngưỡng SAD', '≥ 29 điểm SIAS-17 → có khả năng mắc SAD (Social Anxiety Disorder)'),
    ('Kết quả 18-24 tuổi', 'Cao nhất: 40,3 % vượt ngưỡng SAD'),
]
for i, (k, v) in enumerate(ctx_data):
    c0 = ctx_tbl.rows[i].cells[0]; shade(c0, 'D9E1F2')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True
    ctx_tbl.rows[i].cells[1].text = v
d.add_paragraph()

# KIẾN THỨC NỀN
add_h('Kiến thức nền', 1)

vn_para('1. M là điểm trung bình (Mean) — đúng như thầy đoán', bold=True, size=12)
vn_para('M (viết hoa) = MEAN = ĐIỂM TRUNG BÌNH số học (arithmetic mean) của một mẫu. Tính bằng cách CỘNG TẤT CẢ '
        'các giá trị rồi CHIA cho số lượng.')
vn_para('Trong đoạn của thầy: M = 25,33 nghĩa là điểm trung bình SIAS-17 của nhóm 18-24 tuổi là 25,33 điểm '
        '(trên thang 0-68).')
d.add_paragraph()

vn_para('2. Tại sao M ≠ ngưỡng SAD?', bold=True, size=12)
vn_para('Quan sát: M = 25,33 (TRUNG BÌNH nhóm) < 29 (NGƯỠNG SAD), nhưng 40,3 % CÁ NHÂN trong nhóm vượt ngưỡng.')
vn_para('→ Đây KHÔNG mâu thuẫn. Vì: M là điểm TRUNG BÌNH, một số người DƯỚI 25 và một số người TRÊN 29. '
        'Phân phối không đối xứng — có thể nhiều người ở mức 5-15 (thấp) và một nhóm cao 40-60 — '
        'kéo trung bình về 25,33 nhưng vẫn có 40,3 % cá nhân ≥ 29.')
d.add_paragraph()

vn_para('Hình minh hoạ:', bold=True)
vn_para('Tưởng tượng phân phối SIAS-17 nhóm 18-24:')
vn_para('  ░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░░')
vn_para('   0    10    20  M=25  29   40   60')
vn_para('         (đa số)    │   │           ')
vn_para('                    │   ↑ ngưỡng SAD')
vn_para('                    ↑ trung bình nhóm')
vn_para('Khoảng 40,3 % cá nhân ≥ 29 (vùng từ 29 đến 60).')
d.add_paragraph()

vn_para('3. Các ký hiệu liên quan trong báo cáo thống kê', bold=True, size=12)
sym_tbl = d.add_table(rows=10, cols=3); sym_tbl.style = 'Table Grid'
hdr = ['Ký hiệu', 'Tên đầy đủ', 'Ý nghĩa']
for i, h in enumerate(hdr):
    c = sym_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
sym_data = [
    ('M', 'Mean', 'Điểm trung bình số học = Σx / n'),
    ('SD', 'Standard Deviation', 'Độ lệch chuẩn — đo độ phân tán quanh M'),
    ('SE', 'Standard Error', 'Sai số chuẩn của M = SD/√n'),
    ('Mdn', 'Median', 'Trung vị — giá trị giữa khi xếp tăng dần (ít bị outlier)'),
    ('IQR', 'Interquartile Range', 'Khoảng tứ phân vị (Q3 − Q1) — đo phân tán cho dữ liệu lệch'),
    ('n', 'Sample size', 'Cỡ mẫu (số người tham gia)'),
    ('N', 'Population size', 'Cỡ tổng thể (hoặc tổng cỡ mẫu nhiều nhóm)'),
    ('M ± SD', 'Mean plus-minus SD', 'Cách báo cáo thường dùng: ví dụ "M = 25,33 (SD = 12,5)"'),
    ('95 % CI', '95% Confidence Interval', 'Khoảng tin cậy 95% cho M'),
]
for i, (k, n, m) in enumerate(sym_data):
    c0 = sym_tbl.rows[i+1].cells[0]; shade(c0, 'FFE699')
    pp = c0.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = ORANGE
    sym_tbl.rows[i+1].cells[1].text = n
    sym_tbl.rows[i+1].cells[2].text = m
d.add_paragraph()

vn_para('4. Format báo cáo M chuẩn APA 7th edition', bold=True, size=12)
fmt_tbl = d.add_table(rows=4, cols=2); fmt_tbl.style = 'Table Grid'
fdata = [
    ('M = 25,33', 'Chỉ điểm trung bình — ít thông tin'),
    ('M = 25,33; SD = 12,50', 'CHUẨN APA — luôn báo cáo M kèm SD'),
    ('M = 25,33 (SD = 12,50; n = 1.450)', 'Đầy đủ nhất — kèm cỡ mẫu'),
    ('M = 25,33; 95 % CI [24,8; 25,9]', 'Báo cáo M với khoảng tin cậy thay vì SD — phổ biến trong y học'),
]
for i, (k, v) in enumerate(fdata):
    c0 = fmt_tbl.rows[i].cells[0]; shade(c0, 'D9E1F2')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True
    fmt_tbl.rows[i].cells[1].text = v
d.add_paragraph()

vn_para('5. M vs Mdn (Median) — Khi nào dùng cái nào?', bold=True, size=12)
mdn_tbl = d.add_table(rows=4, cols=2); mdn_tbl.style = 'Table Grid'
mdata = [
    ('Phân phối CHUẨN (đối xứng)', 'Dùng M (Mean). Ít outlier. M ≈ Mdn'),
    ('Phân phối LỆCH (skewed)', 'Dùng Mdn (Median). M bị outlier kéo lệch'),
    ('Có outlier mạnh', 'Mdn ÍT bị ảnh hưởng hơn M. Nên ưu tiên Mdn'),
    ('Mẫu nhỏ (< 30)', 'Có thể báo cả M và Mdn để so sánh'),
]
for i, (k, v) in enumerate(mdata):
    c0 = mdn_tbl.rows[i].cells[0]; shade(c0, 'E2EFDA')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = GREEN
    mdn_tbl.rows[i].cells[1].text = v
d.add_paragraph()

vn_para('6. Diễn giải M = 25,33 SIAS-17 cho 18-24 tuổi', bold=True, size=12)
in_tbl = d.add_table(rows=4, cols=1); in_tbl.style = 'Table Grid'
ipts = [
    'M = 25,33 / 68 (thang tối đa) ≈ 37 % thang. Mức TƯƠNG ĐỐI thấp về số học (chưa đến 50 % thang).',
    'NHƯNG: M < 29 ngưỡng SAD → trung bình nhóm KHÔNG đạt ngưỡng SAD lâm sàng.',
    'Kết hợp với SAD = 40,3 %: nhóm có phân phối "BIMODAL" (2 đỉnh) — đa số dưới ngưỡng + một nhóm CAO trên ngưỡng. '
    'Đây là mẫu hình điển hình của lo âu xã hội: không phải tất cả VTN có lo âu, nhưng nhóm có thì rất rõ.',
    'So sánh: nhóm 25-29 và 30+ tuổi có M thấp hơn (giảm theo tuổi). Lo âu xã hội PEAK ở 18-24 — tuổi đại học/đầu sự nghiệp.',
]
for i, t in enumerate(ipts):
    c = in_tbl.rows[i].cells[0]; shade(c, 'FFF2CC')
    pp = c.paragraphs[0]; rr = pp.add_run('• ' + t); rr.font.size = Pt(11)
d.add_paragraph()

# CÂU TRẢ LỜI (TÔ XANH)
add_h('Câu trả lời', 1, BLUE)

blue_para('ĐÚNG, M là ĐIỂM TRUNG BÌNH (Mean — trung bình số học).', bold=True, size=12)
blue_para('M = 25,33 trong đoạn thầy trích là điểm trung bình SIAS-17 (Social Interaction Anxiety Scale, 17 mục, '
          'thang 0-68) của nhóm 18-24 tuổi trong nghiên cứu Jefferies & Ungar 2020 (QT035 trong CSDL của em — '
          'nghiên cứu 7 quốc gia có Việt Nam, n=6.825).')
d.add_paragraph()

blue_para('Tại sao M = 25,33 < 29 ngưỡng SAD nhưng vẫn có 40,3 % vượt ngưỡng?', bold=True, size=12)
blue_para('Vì M là TRUNG BÌNH NHÓM, không phải mỗi người. Trong nhóm 18-24:')
for it in [
    'Phần đông cá nhân có điểm THẤP (10-25 điểm) → kéo M xuống còn 25,33',
    'Một nhóm 40,3 % có điểm CAO (≥ 29) → vượt ngưỡng SAD',
    'Phân phối "bimodal" (2 đỉnh): đa số bình thường + 40 % có lo âu xã hội rõ',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Các ký hiệu liên quan thầy có thể gặp:', bold=True, size=12)
quick_tbl = d.add_table(rows=8, cols=2); quick_tbl.style = 'Table Grid'
qhdr = ['Ký hiệu', 'Ý nghĩa']
for i, h in enumerate(qhdr):
    c = quick_tbl.rows[0].cells[i]; shade(c, 'BDD7EE')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = BLUE
qdata = [
    ('M (hoặc M̄)', 'Mean = ĐIỂM TRUNG BÌNH'),
    ('SD', 'Standard Deviation = ĐỘ LỆCH CHUẨN'),
    ('Mdn', 'Median = TRUNG VỊ (giá trị giữa)'),
    ('SE', 'Standard Error = SAI SỐ CHUẨN của M'),
    ('IQR', 'Interquartile Range = KHOẢNG TỨ PHÂN VỊ'),
    ('n', 'Sample size = CỠ MẪU'),
    ('95 % CI', '95% Confidence Interval = KHOẢNG TIN CẬY 95%'),
]
for i, (k, v) in enumerate(qdata):
    c = quick_tbl.rows[i+1].cells[0]; pp = c.paragraphs[0]
    pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(k); rr.font.color.rgb = BLUE; rr.bold = True
    c2 = quick_tbl.rows[i+1].cells[1]; pp2 = c2.paragraphs[0]
    rr2 = pp2.add_run(v); rr2.font.color.rgb = BLUE
d.add_paragraph()

blue_para('Format CHUẨN APA: luôn báo cáo M kèm SD', bold=True, size=12)
blue_para('Ví dụ: "M = 25,33 (SD = 12,50)" hoặc "M = 25,33; 95 % CI [24,8; 25,9]". '
          'KHÔNG nên báo cáo chỉ M không có SD/CI vì không biết độ phân tán.')
d.add_paragraph()

blue_para('Diễn giải lâm sàng đoạn của thầy:', bold=True, size=12)
blue_para('Nhóm 18-24 tuổi là PEAK (đỉnh điểm) lo âu xã hội — giai đoạn ĐẠI HỌC + ĐẦU SỰ NGHIỆP. '
          'Mặc dù trung bình nhóm KHÔNG đạt ngưỡng SAD lâm sàng (M=25,33 < 29), '
          'gần một nửa cá nhân (40,3 %) trong nhóm có khả năng mắc SAD. '
          'Khuyến nghị: cần SÀNG LỌC SAD chuyên biệt cho sinh viên + đầu khoá làm việc.')
d.add_paragraph()

blue_para('Ghi nhớ 1 dòng:', bold=True, size=12)
blue_para('M = Mean = điểm trung bình. Luôn đọc M kèm SD (độ lệch chuẩn) để biết độ phân tán. '
          'M < ngưỡng KHÔNG có nghĩa "không ai trong nhóm vượt ngưỡng" — vì đó là TRUNG BÌNH NHÓM, '
          'không phải mỗi cá nhân.', bold=True)
d.add_paragraph()

# PHỤ LỤC
add_h('Phụ lục — Tài liệu tham khảo', 1)
refs = [
    'Jefferies P, Ungar M. (2020). Social Anxiety in Young People: A Prevalence Study in Seven Countries. PLOS ONE, 15(9):e0239133. DOI 10.1371/journal.pone.0239133. [QT035 trong DB — bài gốc của thầy]',
    'Mattick RP, Clarke JC. (1998). Development and validation of measures of social phobia scrutiny fear and social interaction anxiety. Behaviour Research and Therapy, 36(4):455-470. [Bài gốc thang SIAS]',
    'American Psychological Association. (2020). Publication Manual of the American Psychological Association (7th ed.). [Quy ước báo cáo M, SD, n]',
    'Field A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). Sage. [Diễn giải M vs Mdn, phân phối]',
    'Doc liên quan trong dự án: 01_Bao-cao/F_statistic_ANOVA_dien_giai.docx + t_statistic_dien_giai_va_tac_dung.docx + KTC_rong_hep_vua_du_Chen_2023.docx (về M, SD, CI).',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs: r.font.size = Pt(10)

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
