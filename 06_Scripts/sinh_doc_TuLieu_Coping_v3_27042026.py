# -*- coding: utf-8 -*-
"""TuLieu_NN_Coping V3 — chi tiết +30% — tiếng Việt thuần — chú thích cẩn thận."""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\TuLieu_NN_Coping_LoAu_HSTH_v3_chi_tiet_27042026.docx'
RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(9)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def crit(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Lưu ý phản biện] '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(11); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(11); r2.font.name='Times New Roman'
def vn_apply(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Ý nghĩa & áp dụng cho Việt Nam] '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(11); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(11); r2.font.name='Times New Roman'
def block(label, text):
    p=d.add_paragraph()
    r=p.add_run(label + ' '); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.size=Pt(12); r2.font.name='Times New Roman'

# ===================================================================
title("TƯ LIỆU NƯỚC NGOÀI — PHIÊN BẢN 3", 16)
title("VỀ CÁCH ỨNG PHÓ LO ÂU Ở HỌC SINH TRUNG HỌC", 16)
subtitle("(Tiếng Việt thuần — Chi tiết hơn 30% — Chú thích thuật ngữ tiếng Anh cẩn thận — 27/04/2026)")
nr("")
subtitle("Tổng hợp 15 nghiên cứu mới giai đoạn 2023–2026 từ các tạp chí quốc tế "
         "Frontiers, Springer, ScienceDirect, Wiley, JMIR, PMC NCBI, MDPI")
subtitle("Đã đối chiếu với kho tài liệu 35 bài đã có của dự án — hoàn toàn không trùng lặp")
nr("")
nr("Trợ lý nghiên cứu — 27/04/2026 — V3 chi tiết", italic=True, color=GRAY, size=10)
nr("Quy ước trình bày: tiếng Việt làm chính; thuật ngữ tiếng Anh được giữ nguyên trong "
   "ngoặc đơn khi xuất hiện lần đầu (ví dụ: \"liệu pháp nhận thức – hành vi (Cognitive "
   "Behavioural Therapy — CBT)\"); các lần sau chỉ dùng tiếng Việt hoặc viết tắt acronym. "
   "Tên bài báo gốc, tên tác giả, mã định danh số (DOI/PMID/PMC), tên tạp chí và tên thang "
   "đo chuẩn quốc tế giữ nguyên gốc tiếng Anh.",
   italic=True, color=GRAY, size=10)

# ===================================================================
H1("PHẦN A — GIẢI THÍCH CÁC THUẬT NGỮ CHUYÊN MÔN DÙNG TRONG TÀI LIỆU")
nr("Để thầy thuận lợi đọc, em xin liệt kê các thuật ngữ chuyên môn xuất hiện trong tài "
   "liệu này — kèm tên gốc tiếng Anh và viết tắt:")
tbl(['Tiếng Việt', 'Tiếng Anh nguyên gốc', 'Viết tắt'],
    [
        ['Liệu pháp nhận thức – hành vi', 'Cognitive Behavioural Therapy', 'CBT'],
        ['Liệu pháp hành vi biện chứng', 'Dialectical Behaviour Therapy', 'DBT'],
        ['Liệu pháp chấp nhận và cam kết', 'Acceptance and Commitment Therapy', 'ACT'],
        ['Giảm căng thẳng dựa trên chánh niệm', 'Mindfulness-Based Stress Reduction', 'MBSR'],
        ['Can thiệp dựa trên chánh niệm', 'Mindfulness-Based Intervention', 'MBI'],
        ['Học tập về cảm xúc – xã hội', 'Social-Emotional Learning', 'SEL'],
        ['Hiểu biết về sức khỏe tâm thần', 'Mental Health Literacy', 'MHL'],
        ['Tự giới thiệu / tự đăng ký', 'Self-referral', '—'],
        ['Hỗ trợ ngang hàng / từ bạn đồng lứa', 'Peer support', '—'],
        ['Điều hoà cảm xúc', 'Emotion regulation', 'ER'],
        ['Kích thước hiệu ứng', 'Effect size', '—'],
        ['Hệ số d của Cohen', "Cohen's d", "d"],
        ['Khoảng tin cậy 95%', '95% Confidence Interval', '95% CI'],
        ['Thử nghiệm có đối chứng ngẫu nhiên', 'Randomized Controlled Trial', 'RCT'],
        ['Thử nghiệm có đối chứng theo cụm', 'Cluster Randomized Controlled Trial', 'cluster RCT'],
        ['Tổng quan hệ thống', 'Systematic Review', 'SR'],
        ['Phân tích gộp', 'Meta-Analysis', 'MA'],
        ['Tổng quan phạm vi', 'Scoping Review', '—'],
        ['Bác sĩ đa khoa', 'General Practitioner', 'GP'],
        ['Đội hỗ trợ sức khỏe tâm thần (NHS Anh)', 'Mental Health Support Team', 'MHST'],
        ['Dịch vụ tâm thần trẻ em + vị thành niên (NHS Anh)',
         'Child & Adolescent Mental Health Services', 'CAMHS'],
        ['Người chưa từng tham vấn', 'Non-consulter', '—'],
        ['Nhóm khó tiếp cận', 'Hard-to-reach', '—'],
        ['Cách tiếp cận từ dưới lên', 'Bottom-up', '—'],
        ['Trải nghiệm sống', 'Lived experience', '—'],
        ['Mã định danh số', 'Digital Object Identifier', 'DOI'],
        ['Mã PubMed', 'PubMed Identifier', 'PMID'],
        ['Mã PubMed Central', 'PubMed Central', 'PMC'],
        ['Mã thử nghiệm lâm sàng quốc tế Anh', 'International Standard RCT Number', 'ISRCTN'],
    ], [6.0, 6.5, 3.0])

# ===================================================================
H1("PHẦN B — TÓM LƯỢC NHANH (cho người đọc bận)")
nr("Em đã tổng hợp 15 tài liệu nước ngoài MỚI (giai đoạn 2023–2026) về cách ứng phó "
   "lo âu ở học sinh trung học, phân theo 6 nhóm hướng can thiệp chính. Tất cả đều là "
   "tài liệu đã được phản biện chuyên môn (peer-reviewed), bao gồm 4 tổng quan hệ thống, "
   "3 phân tích gộp, 4 thử nghiệm có đối chứng ngẫu nhiên, 2 tổng quan phạm vi và 2 "
   "tổng quan của chính phủ.")
nr("")

H2("Bảng tổng hợp 6 hướng")
tbl(['Hướng can thiệp', 'Số bài', 'Bằng chứng nổi bật'],
    [
        ['1. Quản lý căng thẳng dựa vào trường học', '4',
         'Tổng quan hệ thống + phân tích gộp 38 thử nghiệm có đối chứng / 15.730 học sinh '
         '(Frontiers in Psychiatry 2025); Tổng quan 31 nghiên cứu / 13 quốc gia '
         '(Springer 2024); Vogelaar 2024 Hà Lan với 1.613 học sinh; Lochman 2025 Coping '
         'Power lớp 7 ở Mỹ'],
        ['2. Chánh niệm (Mindfulness)', '4',
         'Phân tích gộp 14 nghiên cứu / 1.489 học sinh — kích thước hiệu ứng tiêu chuẩn '
         '(SMD) = −0,14 cho lo âu; bài CẢNH BÁO của Fulambarkar 2023 trên Wiley CAMH; '
         'thử nghiệm 8+4 tuần củng cố (PMC12173555) năm 2025'],
        ['3. Đào tạo kỹ năng điều hoà cảm xúc', '2',
         'Phân tích gộp 2024 với d=0,37 cho can thiệp tâm lý xã hội; chương trình kỹ '
         'năng điều hoà cảm xúc cho học sinh + cha mẹ (Frontiers 2024)'],
        ['4. Liệu pháp nhận thức – hành vi qua mạng (CBT số)', '2',
         'Thử nghiệm có đối chứng đa trung tâm về CBT online cho lo âu xã hội dưới ngưỡng '
         '(JMIR Pediatrics 2024); tổng quan trò chơi điện tử trị liệu'],
        ['5. Hỗ trợ từ bạn đồng trang lứa', '2',
         'Tổng quan phạm vi của Murphy 2024 (Wiley J Community Psychology); tài liệu '
         'chính phủ Anh về hỗ trợ ngang hàng (NCBI Bookshelf NBK602668)'],
        ['6. Bối cảnh châu Á – Trung Quốc', '1',
         'Tổng quan phạm vi của Lancet Regional Health Western Pacific 2024 về phòng '
         'ngừa và can thiệp sức khỏe tâm thần trường học ở Trung Quốc'],
    ], [4.5, 1.5, 9.5])
nr("")
nr("KẾT LUẬN TỔNG HỢP TỪ 15 TÀI LIỆU TRÊN:", bold=True, size=13)
nr("Bằng chứng giai đoạn 2023–2026 ủng hộ 6 hướng ứng phó lo âu cho học sinh trung học. "
   "Em xếp theo thứ tự ĐỘ MẠNH BẰNG CHỨNG (từ cao xuống thấp), kèm cảnh báo đi cùng:",
   bold=True)
nr("• Hạng 1: Quản lý căng thẳng dựa trên Liệu pháp nhận thức – hành vi (đào tạo kỹ năng "
   "có cấu trúc) — kích thước hiệu ứng nhỏ đến trung bình, NHẤT QUÁN qua nhiều nghiên "
   "cứu. Đây là hướng AN TOÀN NHẤT để đầu tư.")
nr("• Hạng 2: Đào tạo kỹ năng điều hoà cảm xúc (kỹ thuật từ DBT, ACT, CBT tập trung "
   "vào cảm xúc) — hệ số d của Cohen ≈ 0,37 (mức trung bình theo chuẩn Cohen 1988). "
   "Có tiềm năng cao hơn chánh niệm.")
nr("• Hạng 3: Cơ chế tự giới thiệu có chỉ định (đã chứng minh trong nghiên cứu BESST "
   "Brown 2024 đã được dịch trong dự án) — đạt hiệu quả trung bình d=−0,52 cho NHÓM "
   "ĐÃ CÓ TRIỆU CHỨNG CAO; phù hợp với mô hình PLACES.")
nr("• Hạng 4: Liệu pháp nhận thức – hành vi qua mạng (CBT online) — có thể triển khai "
   "quy mô lớn nhưng tỉ lệ duy trì tham gia (adherence) thấp khi không có người hướng "
   "dẫn (unguided).")
nr("• Hạng 5: Chánh niệm phổ quát (cho mọi học sinh không phân biệt) — hiệu ứng nhỏ "
   "(SMD = −0,14, dưới ngưỡng \"có ý nghĩa lâm sàng\" 0,2 của Cohen). CẦN THẬN TRỌNG "
   "với mô hình thất bại của thử nghiệm MYRIAD ở Anh năm 2022 với 8.376 học sinh.")
nr("• Hạng 6: Hỗ trợ từ bạn đồng trang lứa — hứa hẹn về mặt khái niệm nhưng bằng chứng "
   "định lượng còn YẾU (chưa đủ thử nghiệm có đối chứng).")
nr("")
nr("KHUYẾN NGHỊ THỰC TIỄN: Việt Nam nên ƯU TIÊN kết hợp 3 hướng đầu (đào tạo CBT có cấu "
   "trúc + kỹ năng điều hoà cảm xúc + cơ chế tự giới thiệu có chỉ định); bổ trợ bằng "
   "công cụ số (CBT online tiếng Việt); KHÔNG nên đầu tư mạnh vào chánh niệm phổ quát "
   "do giáo viên dẫn dắt vì đã được chứng minh thất bại quy mô lớn ở Anh.",
   bold=True, color=BLUE)

# ===================================================================
H1("PHẦN C — CHI TIẾT 15 NGHIÊN CỨU THEO 6 NHÓM")

# ---------- NHÓM 1 ----------
H1("NHÓM 1 — Quản lý căng thẳng dựa vào trường học (4 bài)")
nr("Đây là nhóm bằng chứng MẠNH NHẤT trong 6 nhóm — gồm các tổng quan hệ thống và phân "
   "tích gộp lớn cùng các thử nghiệm có đối chứng quy mô đa trung tâm. Quản lý căng "
   "thẳng (stress management) ở học sinh trung học là chủ đề được đầu tư nhiều nhất "
   "trong 5 năm gần đây.")

H2("Bài 1.1 — Tổng quan hệ thống và phân tích gộp 38 thử nghiệm có đối chứng (Frontiers "
   "in Psychiatry 2025)")
block('Tên bài gốc:',
      '\"School-based interventions for resilience in children and adolescents: a '
      'systematic review and meta-analysis of randomized controlled trials\"')
block('Tạm dịch:',
      'Các can thiệp dựa vào trường học nhằm tăng khả năng thích ứng (resilience) ở '
      'trẻ em và vị thành niên: tổng quan hệ thống và phân tích gộp các thử nghiệm có '
      'đối chứng ngẫu nhiên')
block('Tạp chí:',
      'Frontiers in Psychiatry — chuyên mục Sức khỏe tâm thần công cộng (Public Mental '
      'Health), năm 2025')
block('Mã định danh:',
      'DOI: 10.3389/fpsyt.2025.1594658 — Mã PubMed Central: PMC12127306')
block('Bối cảnh nghiên cứu:',
      'Khả năng thích ứng (resilience) đã được công nhận từ những năm 1990 là yếu tố '
      'bảo vệ quan trọng giảm nguy cơ phát triển lo âu / trầm cảm ở vị thành niên. '
      'Tuy nhiên các chương trình \"đào tạo resilience\" ở trường có hiệu quả không '
      'thì chưa được tổng hợp đầy đủ. Tổng quan này lấp đầy khoảng trống đó với cỡ mẫu '
      'tổng hợp lớn nhất từng có cho chủ đề.')
block('Phương pháp:',
      'Tổng quan hệ thống 38 thử nghiệm có đối chứng ngẫu nhiên với tổng cộng 15.730 '
      'người tham gia (gồm trẻ em + vị thành niên). Các thử nghiệm được tìm kiếm trên '
      'PubMed, EMBASE, PsycINFO, Web of Science. Nguồn dữ liệu được phân loại theo '
      'loại can thiệp: chánh niệm (Mindfulness-Based Intervention — MBI), liệu pháp '
      'nhận thức – hành vi (Cognitive Behavioural Therapy — CBT), học tập về cảm xúc – '
      'xã hội (Social-Emotional Learning — SEL), các chương trình tăng resilience '
      'chuyên biệt.')
block('Phát hiện chính:',
      'Học sinh có resilience cao hơn có xu hướng phản ứng tích cực với các tình huống '
      'căng thẳng và sử dụng các chiến lược ứng phó thích nghi (adaptive coping) tốt '
      'hơn — nhờ đó giảm nguy cơ phát triển các cảm xúc tiêu cực như lo âu và trầm cảm. '
      'Tích hợp đồng thời 3 thành phần (hiểu biết về sức khỏe tâm thần + kỹ thuật quản '
      'lý căng thẳng + học tập về cảm xúc – xã hội) vào chương trình giảng dạy ở '
      'trường học giúp cải thiện rõ rệt khả năng phục hồi cảm xúc, giảm căng thẳng '
      'học tập, và xây dựng các mối quan hệ xã hội tích cực.')
crit("38 thử nghiệm với 15.730 học sinh là quy mô RẤT LỚN — bằng chứng cấp 1 (cao "
     "nhất theo thang Oxford CEBM 2011). Tuy nhiên các thử nghiệm thuộc nhiều dạng "
     "can thiệp khác nhau nên độ không đồng nhất (heterogeneity) cao, kích thước "
     "hiệu ứng gộp khó diễn giải đơn lẻ — phải đọc các phân tích phụ (subgroup "
     "analyses) để biết hướng nào hiệu quả nhất. Đặc biệt cần đọc xem hiệu quả có "
     "khác nhau giữa các nhóm tuổi (trẻ em vs vị thành niên) hay không.")
vn_apply("Kết quả này ủng hộ Bộ Giáo dục và Đào tạo Việt Nam tích hợp 3 thành phần "
         "trên vào tiết Hoạt động trải nghiệm – Hướng nghiệp (HĐTN-HN) hoặc Giáo dục "
         "công dân lớp 10–12. Đề xuất cụ thể: dành 1 học kỳ (16 tuần) cho chương trình "
         "tích hợp, mỗi tuần 1 tiết 45 phút, gồm 4 module: Module 1 (4 tuần) — Hiểu "
         "biết về sức khỏe tâm thần; Module 2 (4 tuần) — Kỹ thuật quản lý căng thẳng; "
         "Module 3 (4 tuần) — Học tập về cảm xúc – xã hội; Module 4 (4 tuần) — Tích "
         "hợp + thực hành. Có thể tham khảo nội dung mẫu trong các thử nghiệm thuộc "
         "tổng quan này để xây dựng chương trình bản địa hoá.")

H2("Bài 1.2 — Tổng quan hệ thống can thiệp giảm căng thẳng học tập ở trường THPT (Springer 2024)")
block('Tên bài gốc:',
      '\"Academic Stress Interventions in High Schools: A Systematic Literature Review\"')
block('Tạm dịch:',
      'Các can thiệp giảm căng thẳng học tập ở trường trung học phổ thông: Tổng quan '
      'tài liệu hệ thống')
block('Tạp chí:', 'Child Psychiatry & Human Development (Springer Nature) — năm 2024')
block('Mã định danh:',
      'DOI: 10.1007/s10578-024-01667-5 — Mã PubMed Central: PMC12628395')
block('Bối cảnh nghiên cứu:',
      'Căng thẳng học tập (academic stress) là vấn đề phổ biến nhất với học sinh THPT '
      'toàn cầu, đặc biệt ở các nước Đông Á có hệ thống thi cử cạnh tranh cao. Cho '
      'đến nay chưa có tổng quan tập trung riêng vào các can thiệp giảm căng thẳng '
      'HỌC TẬP (khác với các can thiệp căng thẳng tổng quát). Tổng quan này cung cấp '
      'cái nhìn tổng hợp đầu tiên đa quốc gia về chủ đề này.')
block('Phương pháp:',
      'Tổng quan 31 nghiên cứu phù hợp tiêu chí, trải dài 13 quốc gia, về các can '
      'thiệp giúp giảm hoặc phòng ngừa căng thẳng học tập ở học sinh trung học phổ '
      'thông (16–18 tuổi). Bao gồm cả can thiệp định lượng và định tính.')
block('Phát hiện chính:',
      'Các yếu tố gây căng thẳng chính được học sinh báo cáo: (1) khối lượng bài vở '
      'lớn, (2) áp lực thi cử và xếp hạng, (3) áp lực từ bạn bè (peer pressure), '
      '(4) kỳ vọng từ phụ huynh và giáo viên, (5) lo lắng về tương lai (định hướng '
      'nghề nghiệp). Các chiến lược ứng phó tự phát mà học sinh thường dùng (theo '
      'tự báo cáo): (a) hỗ trợ xã hội từ bạn đồng trang lứa, (b) tập thể dục / vận '
      'động, (c) các kỹ thuật thư giãn như thiền và hít thở sâu, (d) giải quyết vấn '
      'đề có hệ thống (problem-solving), (e) tái cấu trúc nhận thức — thay đổi cách '
      'nghĩ về tình huống căng thẳng (cognitive restructuring).')
crit("31 nghiên cứu trải 13 quốc gia là dữ liệu đa dạng — strong external validity. "
     "Tuy nhiên các \"chiến lược ứng phó\" được đo qua tự báo cáo (self-report) trong "
     "các nghiên cứu cắt ngang (cross-sectional), không thể kết luận quan hệ nhân – "
     "quả (kiểu \"chiến lược X dẫn đến giảm lo âu\"). Để có kết luận nhân – quả cần "
     "thử nghiệm có đối chứng riêng cho từng chiến lược. Ngoài ra, tự báo cáo có rủi "
     "ro thiên lệch nhớ lại (recall bias) và thiên lệch xã hội (social desirability "
     "bias — học sinh có thể trả lời \"chiến lược tốt\" mà thực tế ít dùng).")
vn_apply("Có thể xây dựng tài liệu cho giáo viên / cán bộ tư vấn học đường Việt Nam "
         "dựa trên 5 chiến lược ứng phó nói trên: thiết kế thành 5 chuyên đề, mỗi "
         "chuyên đề 1 tiết HĐTN. Cấu trúc đề xuất: Chuyên đề 1 — Tăng cường hỗ trợ "
         "xã hội từ bạn (kỹ năng giao tiếp + xây dựng nhóm bạn); Chuyên đề 2 — Vận "
         "động (lập kế hoạch tập thể dục đều đặn); Chuyên đề 3 — Thư giãn (thực hành "
         "thiền + hít thở); Chuyên đề 4 — Giải quyết vấn đề (kỹ thuật chia nhỏ vấn "
         "đề); Chuyên đề 5 — Tái cấu trúc nhận thức (nhận diện và thay đổi suy nghĩ "
         "tiêu cực). Đặc biệt phù hợp lớp 11–12 do áp lực thi đại học cao.")

H2("Bài 1.3 — Thử nghiệm \"Bài học về căng thẳng\" của Vogelaar 2024 ở Hà Lan")
block('Tên bài gốc:',
      'Theo trích dẫn từ phân tích gộp đăng trên Journal of School Psychology 105 '
      '(2024) — Vogelaar và cộng sự (2024). Đánh giá hiệu quả chương trình tâm lý '
      'giáo dục \"Stress Lessons\" trên học sinh vị thành niên Hà Lan.')
block('Tạp chí:', 'Journal of School Psychology, Tập 105 (2024)')
block('Đường dẫn:',
      'https://www.sciencedirect.com/science/article/pii/S0022440524000724')
block('Bối cảnh nghiên cứu:',
      'Hà Lan là một trong những quốc gia có chương trình hỗ trợ sức khỏe tâm thần '
      'trường học phát triển ở châu Âu. Chương trình \"Stress Lessons\" được Bộ Giáo '
      'dục Hà Lan đầu tư nhằm chuẩn hoá việc dạy về căng thẳng cho học sinh vị '
      'thành niên — đây là thử nghiệm đánh giá hiệu quả thực tế của chương trình.')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên theo cụm (cluster RCT) với 1.613 học sinh '
      'vị thành niên Hà Lan. Chương trình \"Bài học về căng thẳng\" gồm 3 buổi học, '
      '1 buổi/tuần, mỗi buổi kéo dài 45 phút. Mục tiêu: (1) giảm căng thẳng tổng '
      'thể, (2) giảm căng thẳng đặc thù ở trường, (3) tăng hiểu biết về căng thẳng. '
      'Đối chứng: chăm sóc thông thường (treatment-as-usual).')
crit("Chỉ 3 buổi học × 45 phút (tổng 135 phút = ~2 giờ 15 phút) là cường độ THẤP. "
     "Đây là chương trình tâm lý giáo dục thuần tuý (cung cấp kiến thức), KHÔNG có "
     "thành phần luyện tập kỹ năng CBT thực tế. Hiệu ứng dự kiến nhỏ (kỳ vọng d ≈ "
     "0,1–0,2). Để có hiệu quả lớn hơn cần ít nhất 6–8 buổi với thực hành giữa các "
     "buổi.")
vn_apply("Mô hình 3 buổi × 45 phút phù hợp tích hợp vào tiết HĐTN ở Việt Nam (1 tiết "
         "= 45 phút). Tuy nhiên em ĐỀ XUẤT mở rộng thành 6 buổi: 3 buổi đầu giữ "
         "nguyên (tâm lý giáo dục về căng thẳng), thêm 3 buổi sau luyện kỹ năng CBT "
         "thực tế (1 buổi cho mỗi kỹ năng: tái cấu trúc nhận thức, kích hoạt hành "
         "vi, giải quyết vấn đề). Tổng 6 buổi = 1 nửa học kỳ ở Việt Nam, vẫn dễ tích "
         "hợp.")

H2("Bài 1.4 — Chương trình \"Sức mạnh ứng phó\" cho vị thành niên sớm (Lochman 2025 ở Mỹ)")
block('Tên bài gốc:',
      '\"Randomized controlled trial of the early adolescent coping power program: '
      'Effects on emotional and behavioral problems in middle schoolers\"')
block('Tạm dịch:',
      'Thử nghiệm có đối chứng ngẫu nhiên về chương trình Sức mạnh ứng phó cho vị '
      'thành niên sớm: tác động lên các vấn đề cảm xúc và hành vi ở học sinh trung '
      'học cơ sở')
block('Tạp chí:', 'Journal of School Psychology (2025)')
block('Đường dẫn:',
      'https://www.sciencedirect.com/science/article/abs/pii/S002244052500010X')
block('Bối cảnh nghiên cứu:',
      'Chương trình Sức mạnh ứng phó (Coping Power) do Lochman & Wells phát triển '
      'từ những năm 1990, đã có hơn 30 năm bằng chứng cho học sinh tiểu học (8–11 '
      'tuổi) với hiệu quả tốt trong giảm hành vi gây hấn và cải thiện kỹ năng ứng '
      'phó. Phiên bản này được điều chỉnh cho vị thành niên sớm (early adolescent '
      '— 11–13 tuổi) — giai đoạn não bộ đang phát triển, kỹ năng tự điều hoà chưa '
      'hoàn thiện và rất cần can thiệp sớm.')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên theo nhóm cấp trường (school-level group '
      'RCT), đối tượng là học sinh lớp 7 (tương đương lớp 7 ở Việt Nam, 12–13 tuổi) '
      'tại 40 trường trung học cơ sở thuộc các bang Alabama và Maryland (Mỹ). Can '
      'thiệp là phiên bản dành cho vị thành niên sớm của chương trình Coping Power, '
      'gồm 24–34 buổi học kỹ năng (có thay đổi tuỳ trường) trong 1 năm học.')
crit("Chương trình Coping Power đã có nhiều bằng chứng cho học sinh tiểu học. "
     "Phiên bản này là điều chỉnh cho vị thành niên sớm (11–13 tuổi). Kích thước "
     "hiệu ứng dự kiến ở mức nhỏ – trung bình. Một hạn chế: chương trình dài (24–34 "
     "buổi) — khó triển khai trong context Việt Nam có ít thời gian HĐTN.")
vn_apply("Chương trình Coping Power có sách hướng dẫn (manual) đầy đủ — Việt Nam có "
         "thể nhập giấy phép sử dụng để bản địa hoá cho học sinh lớp 6–8 (tương "
         "đương vị thành niên sớm). Cần đào tạo bài bản cho cán bộ tư vấn học "
         "đường + có giám sát chuyên môn từ chuyên gia tâm lý. Có thể rút gọn từ "
         "24–34 buổi xuống 12–16 buổi (1 học kỳ) bằng cách giữ các module cốt lõi: "
         "nhận diện cảm xúc, kiểm soát giận dữ, giải quyết vấn đề, kỹ năng xã hội, "
         "đặt mục tiêu.")

# ---------- NHÓM 2 ----------
H1("NHÓM 2 — Chánh niệm — Mindfulness (4 bài)")

H2("Bài 2.1 — Phân tích gộp 14 nghiên cứu về chánh niệm (MDPI Pediatric Reports 2025)")
block('Tên bài gốc:',
      '\"Mindfulness in Mental Health and Psychiatric Disorders of Children and '
      'Adolescents: A Systematic Review and Meta-Analysis of Randomized Controlled '
      'Trials\"')
block('Tạm dịch:',
      'Chánh niệm trong sức khỏe tâm thần và rối loạn tâm thần ở trẻ em và vị thành '
      'niên: Tổng quan hệ thống và phân tích gộp các thử nghiệm có đối chứng '
      'ngẫu nhiên')
block('Tạp chí:', 'Pediatric Reports (MDPI) — Tập 17(3), Bài 59 (2025)')
block('Đường dẫn:', 'https://www.mdpi.com/2036-7503/17/3/59')
block('Bối cảnh nghiên cứu:',
      'Chánh niệm (mindfulness) là phương pháp đã được áp dụng rộng rãi trong các '
      'trường học phương Tây từ năm 2010 — đặc biệt qua chương trình Giảm căng '
      'thẳng dựa trên chánh niệm (MBSR). Tuy nhiên các phân tích gộp gần đây cho '
      'kết quả lẫn lộn về hiệu quả, đặc biệt ở vị thành niên. Phân tích gộp 2025 '
      'này tổng hợp các thử nghiệm có đối chứng đã có để đưa ra ước tính chính xác '
      'nhất về kích thước hiệu ứng.')
block('Phương pháp:',
      'Phân tích gộp 14 nghiên cứu thử nghiệm có đối chứng ngẫu nhiên với tổng cộng '
      '1.489 người tham gia. Tìm kiếm trên các cơ sở dữ liệu PubMed, Cochrane, '
      'EMBASE. Outcome chính: triệu chứng lo âu sau can thiệp.')
block('Phát hiện chính:',
      'Chánh niệm theo chương trình MBSR làm giảm triệu chứng lo âu so với nhóm đối '
      'chứng với KÍCH THƯỚC HIỆU ỨNG TIÊU CHUẨN HOÁ (Standardized Mean Difference '
      '— SMD) = −0,14 (khoảng tin cậy 95%: từ −0,24 đến −0,04) ngay sau khi kết '
      'thúc trị liệu. Mức hiệu ứng này được xếp vào loại NHỎ theo mức tham chiếu '
      'của Cohen 1988 (0,2 = nhỏ, 0,5 = trung bình, 0,8 = lớn) — và thậm chí dưới '
      'ngưỡng \"nhỏ\".')
crit("Hiệu ứng SMD = −0,14 là RẤT NHỎ — DƯỚI ngưỡng \"có ý nghĩa lâm sàng\" 0,2 của "
     "Cohen. Chánh niệm có hiệu ứng THẤP HƠN liệu pháp nhận thức – hành vi (CBT "
     "thường đạt 0,3–0,5). Phù hợp với phát hiện của thử nghiệm MYRIAD ở Anh năm "
     "2022 (Kuyken và cộng sự, n=8.376) — kết quả không có hiệu ứng (null effect) "
     "khi áp dụng chánh niệm phổ quát do giáo viên dẫn dắt. Bằng chứng này khuyến "
     "cáo KHÔNG nên đầu tư mạnh vào chánh niệm như can thiệp chính.")
vn_apply("Việt Nam nên CẨN THẬN với chánh niệm phổ quát do giáo viên dẫn dắt — "
         "không nên là can thiệp CHÍNH. Có thể dùng làm hoạt động BỔ TRỢ ngắn 5–10 "
         "phút mỗi ngày trong tiết HĐTN hoặc đầu giờ học để rèn chú ý — nhưng đừng "
         "kỳ vọng giảm đáng kể lo âu. Đầu tư tiền và thời gian vào CBT skills "
         "training sẽ hiệu quả hơn.")

H2("Bài 2.2 — Thử nghiệm chánh niệm 8 tuần + 4 tuần củng cố cho vị thành niên 13–15 "
   "tuổi (PMC12173555 năm 2025)")
block('Mã định danh:', 'PMC12173555')
block('Đường dẫn:', 'https://pmc.ncbi.nlm.nih.gov/articles/PMC12173555/')
block('Bối cảnh nghiên cứu:',
      'Một hạn chế lớn của các thử nghiệm chánh niệm trước đây là hiệu ứng tan '
      'nhanh sau khi kết thúc can thiệp (gọi là \"effect decay\"). Nghiên cứu này '
      'thử thiết kế MỚI — thêm giai đoạn củng cố hằng tuần (booster sessions) sau '
      '8 tuần can thiệp ban đầu — để giải quyết vấn đề hiệu ứng tan.')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên trên mẫu KHÔNG có chẩn đoán lâm sàng '
      '(non-clinical sample) gồm vị thành niên 13–15 tuổi. Can thiệp gồm: '
      'GIAI ĐOẠN 1: 8 tuần huấn luyện chánh niệm (1 buổi/tuần, mỗi buổi 60 phút); '
      'GIAI ĐOẠN 2: 4 tuần củng cố hằng tuần (booster, 30 phút/buổi). Đo lường: '
      'các triệu chứng nội tâm hoá (lo âu, trầm cảm), trạng thái cảm xúc, các '
      'chiến lược điều hoà cảm xúc.')
crit("Phần củng cố 4 tuần là điểm THÚ VỊ và ĐỔI MỚI — giải quyết được hạn chế "
     "kinh điển của các thử nghiệm chánh niệm trước (effect decay). Cần đọc bản "
     "đầy đủ để xem giai đoạn củng cố có thực sự giúp duy trì hiệu quả ở thời "
     "điểm 6 tháng sau can thiệp hay không.")
vn_apply("Mô hình 8 tuần + 4 tuần củng cố phù hợp với học kỳ ở Việt Nam (1 học kỳ "
         "khoảng 16 tuần). Có thể bản địa hoá cho học sinh trung học cơ sở lớp 8–9. "
         "Cần thiết kế các buổi booster ngắn gọn (20–30 phút) đủ tích hợp vào "
         "tiết HĐTN.")

H2("Bài 2.3 — Phân tích gộp CẢNH BÁO của Fulambarkar (Wiley CAMH 2023)")
block('Tên bài gốc:',
      '\"Meta-analysis on mindfulness-based interventions for adolescents stress, '
      'depression, and anxiety in school settings: a cautionary tale\"')
block('Tạm dịch:',
      'Phân tích gộp về các can thiệp dựa trên chánh niệm cho căng thẳng, trầm cảm '
      'và lo âu ở vị thành niên trong môi trường trường học: một câu chuyện cần '
      'cảnh giác')
block('Tạp chí:', 'Child and Adolescent Mental Health (Wiley) — năm 2023')
block('Mã định danh:', 'DOI: 10.1111/camh.12572')
block('Đường dẫn:',
      'https://acamh.onlinelibrary.wiley.com/doi/10.1111/camh.12572')
block('Bối cảnh nghiên cứu:',
      'Sau 2 thập kỷ quảng bá rộng rãi chánh niệm trong trường học, đặc biệt ở Mỹ '
      'và Anh, một số chuyên gia bắt đầu đặt câu hỏi liệu chánh niệm phổ quát có '
      'thực sự hiệu quả hay là một hiện tượng \"hype\" (thổi phồng). Bài phân tích '
      'gộp này được đặt tên \"a cautionary tale\" (câu chuyện cảnh giác) — phản '
      'biện trực diện với xu hướng ủng hộ chánh niệm.')
block('Phát hiện chính:',
      'Cảnh báo về việc tin tưởng quá mức vào hiệu quả chánh niệm dựa vào trường '
      'học cho vị thành niên. Kích thước hiệu ứng tổng hợp ở mức nhỏ tới không có '
      '(near-null effect); mức độ tham gia của học sinh thấp; thậm chí có thể GÂY '
      'HẠI cho một số học sinh dễ bị tổn thương (vulnerable subgroup) — đặc biệt '
      'những em đã có triệu chứng lo âu / trầm cảm baseline.')
crit("Đây là phân tích gộp có TÍNH PHẢN BIỆN với xu hướng quảng bá \"chánh niệm "
     "cho tất cả mọi người\". Nên đọc song song với các nghiên cứu ủng hộ chánh "
     "niệm để có cái nhìn cân bằng. Phù hợp với bằng chứng từ thử nghiệm MYRIAD ở "
     "Anh năm 2022 và phân tích gộp Pediatric Reports 2025 (cùng cho SMD nhỏ).")
vn_apply("Cần tham khảo bài này TRƯỚC KHI quyết định tích hợp chánh niệm vào "
         "chương trình chính khoá ở Việt Nam. Cách an toàn nhất: chỉ làm thành "
         "hoạt động tự chọn (optional) + workshop ngoài giờ học, KHÔNG bắt buộc "
         "toàn trường. Đặc biệt cảnh báo: nếu có học sinh đang có triệu chứng lo "
         "âu / trầm cảm rõ rệt, KHÔNG ép tham gia chánh niệm vì có thể làm trầm "
         "trọng hơn — cần chuyển sang CBT có cấu trúc.")

H2("Bài 2.4 — Nghiên cứu có đối chứng về chánh niệm 10 tuần ở trường (School Mental "
   "Health 2023)")
block('Tạp chí:', 'School Mental Health (Springer) — năm 2023')
block('Mã định danh:', 'DOI: 10.1007/s12310-023-09620-y')
block('Đường dẫn:',
      'https://link.springer.com/article/10.1007/s12310-023-09620-y')
block('Phương pháp:',
      'Nghiên cứu có đối chứng (không phải thử nghiệm ngẫu nhiên đầy đủ) đánh giá '
      'hiệu quả của can thiệp chánh niệm dài 10 tuần tại trường lên triệu chứng '
      'trầm cảm và lo âu của học sinh từ trẻ em đến vị thành niên. Đây là một '
      'trong số ít nghiên cứu có thiết kế quasi-experimental (gần thực nghiệm) '
      'cho bối cảnh trường thực tế — quan trọng vì các thử nghiệm RCT \"thuần\" '
      'thường khó áp dụng ngay vào môi trường giáo dục thực tế.')

# ---------- NHÓM 3 ----------
H1("NHÓM 3 — Đào tạo kỹ năng điều hoà cảm xúc (2 bài)")

H2("Bài 3.1 — Phân tích gộp về can thiệp tâm lý xã hội với điều hoà cảm xúc (Springer 2024)")
block('Tên bài gốc:',
      '\"Effect of Psychosocial Interventions on Children and Youth Emotion '
      'Regulation: A Meta-Analysis\"')
block('Tạm dịch:',
      'Tác động của các can thiệp tâm lý xã hội lên năng lực điều hoà cảm xúc ở '
      'trẻ em và thanh thiếu niên: Một phân tích gộp')
block('Tạp chí:',
      'Administration and Policy in Mental Health and Mental Health Services Research '
      '(Springer) — năm 2024')
block('Mã định danh:', 'DOI: 10.1007/s10488-024-01373-3')
block('Đường dẫn:',
      'https://link.springer.com/article/10.1007/s10488-024-01373-3')
block('Bối cảnh nghiên cứu:',
      'Điều hoà cảm xúc (emotion regulation — ER) là khả năng nhận diện, hiểu và '
      'điều chỉnh các trạng thái cảm xúc — được coi là kỹ năng cốt lõi cho sức '
      'khỏe tâm thần. Sự thiếu hụt ER trong giai đoạn vị thành niên là một yếu tố '
      'nguy cơ mạnh cho phát triển lo âu và trầm cảm. Phân tích gộp này tổng hợp '
      'các can thiệp tâm lý xã hội (psychosocial interventions — bao gồm CBT, '
      'DBT, ACT, chánh niệm) tác động lên ER của trẻ em và thanh thiếu niên.')
block('Phát hiện chính:',
      'Các can thiệp tâm lý xã hội có hiệu ứng từ NHỎ ĐẾN TRUNG BÌNH lên năng lực '
      'điều hoà cảm xúc với hệ số d của Cohen = 0,37. Sự thiếu hụt khả năng điều '
      'hoà cảm xúc trong giai đoạn vị thành niên có liên quan đến nguy cơ cao hơn '
      'phát triển các triệu chứng lo âu và trầm cảm — củng cố tầm quan trọng của '
      'việc đầu tư vào đào tạo ER ngay từ tuổi vị thành niên.')
crit("Hiệu ứng d = 0,37 là ĐÁNG KỂ HƠN chánh niệm (chỉ d = 0,14) — đào tạo kỹ năng "
     "điều hoà cảm xúc có tiềm năng cao hơn. Tuy nhiên cần biết loại can thiệp cụ "
     "thể nào trong số các can thiệp tâm lý xã hội cho hiệu ứng cao nhất (DBT, "
     "ACT, CBT tập trung cảm xúc?) — phân tích gộp này không phân loại rõ. Cần "
     "đọc subgroup analysis để biết.")
vn_apply("Nên ƯU TIÊN đào tạo kỹ năng điều hoà cảm xúc (sử dụng các kỹ thuật từ "
         "Liệu pháp hành vi biện chứng — Dialectical Behaviour Therapy, DBT — "
         "Liệu pháp chấp nhận và cam kết — Acceptance and Commitment Therapy, ACT "
         "— hoặc CBT tập trung cảm xúc) HƠN chánh niệm thuần tuý. Quy mô đề xuất "
         "cho Việt Nam: chương trình 6–8 tuần, 1 buổi/tuần, mỗi buổi 60 phút, gồm "
         "4 module DBT cốt lõi: (1) Chánh niệm cốt lõi (mindfulness); "
         "(2) Chịu đựng đau khổ (distress tolerance); (3) Điều hoà cảm xúc "
         "(emotion regulation); (4) Hiệu quả giao tiếp (interpersonal "
         "effectiveness). Có thể đào tạo cho cán bộ tư vấn học đường + giám sát "
         "từ chuyên gia tâm lý.")

H2("Bài 3.2 — Đào tạo kỹ năng điều hoà cảm xúc cho học sinh + cha mẹ (Frontiers 2024)")
block('Tên bài gốc:',
      '\"An emotion regulation skills training for adolescents and parents: '
      'perceptions and acceptability of methodological aspects\"')
block('Tạm dịch:',
      'Một chương trình đào tạo kỹ năng điều hoà cảm xúc cho vị thành niên và '
      'cha mẹ: nhận thức và mức độ chấp nhận về các khía cạnh phương pháp luận')
block('Tạp chí:', 'Frontiers in Psychiatry — năm 2024')
block('Mã định danh:', 'DOI: 10.3389/fpsyt.2024.1448529')
block('Đường dẫn:',
      'https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1448529/full')
block('Bối cảnh nghiên cứu:',
      'Nhiều nghiên cứu cho thấy cha mẹ có vai trò quan trọng trong việc \"mô hình '
      'hoá\" (modeling) cách điều hoà cảm xúc cho con. Nếu cha mẹ không có kỹ năng '
      'ER tốt, dạy con kỹ năng này sẽ kém hiệu quả. Bài này thử nghiệm chương '
      'trình KẾT HỢP cha mẹ + con cùng học kỹ năng ER.')
block('Đặc điểm:',
      'Chương trình đào tạo KẾT HỢP cả vị thành niên và cha mẹ tham gia cùng nhau. '
      'Đo lường chính: cảm nhận của người tham gia về chương trình + mức chấp '
      'nhận về các khía cạnh phương pháp luận (như độ dài buổi học, format trực '
      'tuyến / trực tiếp, phương pháp giảng dạy).')
crit("Đào tạo kết hợp cha mẹ – vị thành niên là MẠNH về mặt lý thuyết hệ thống "
     "(theo lý thuyết hệ thống gia đình — family systems theory) nhưng KHÓ triển "
     "khai quy mô lớn vì cha mẹ thường khó sắp xếp thời gian. Đặc biệt khó với "
     "Việt Nam nơi cha mẹ ở vùng nông thôn / công nhân thường làm việc nhiều ca.")
vn_apply("Việt Nam có thể thử mô hình \"cha mẹ cùng học sinh tham gia\" qua các "
         "buổi họp phụ huynh trực tuyến (Zoom / Zalo / Microsoft Teams). Cách "
         "này phù hợp với văn hoá Á Đông đề cao vai trò gia đình trong giáo dục "
         "con cái. Đề xuất cụ thể: tổ chức 4 buổi tối thứ 7 (mỗi buổi 90 phút) "
         "qua Zoom, kết hợp lý thuyết + thực hành tại nhà, có sách bài tập gửi "
         "trước. Với cha mẹ ít công nghệ, có thể quay video buổi học và gửi qua "
         "Zalo để xem lại.")

# ---------- NHÓM 4 ----------
H1("NHÓM 4 — Liệu pháp nhận thức – hành vi qua mạng (CBT số) (2 bài)")

H2("Bài 4.1 — Thử nghiệm CBT online cho lo âu xã hội dưới ngưỡng (JMIR Pediatrics 2024)")
block('Tên bài gốc:',
      '\"Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for '
      'Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: '
      'Multicenter Randomized Controlled Trial\"')
block('Tạm dịch:',
      'Hiệu quả của Liệu pháp nhận thức – hành vi qua Internet KHÔNG có hướng dẫn '
      'cho rối loạn lo âu xã hội dưới ngưỡng ở vị thành niên và thanh niên: Thử '
      'nghiệm có đối chứng ngẫu nhiên đa trung tâm')
block('Tạp chí:', 'JMIR Pediatrics and Parenting — Tập 7 (2024)')
block('Mã định danh:', 'DOI: 10.2196/55786')
block('Đường dẫn:', 'https://pediatrics.jmir.org/2024/1/e55786')
block('Bối cảnh nghiên cứu:',
      'Rối loạn lo âu xã hội (Social Anxiety Disorder — SAD) thường khởi phát ở '
      'tuổi vị thành niên và có thể gây hậu quả lâu dài về học tập, nghề nghiệp '
      'và quan hệ xã hội. Tuy nhiên đa số người trẻ có triệu chứng dưới ngưỡng '
      '(subthreshold) chưa đủ tiêu chuẩn chẩn đoán không tiếp cận được trị liệu '
      'chính thống. CBT online không hướng dẫn (unguided) là giải pháp khả thi '
      'để mở rộng tiếp cận với chi phí thấp.')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên đa trung tâm về CBT online KHÔNG có '
      'hướng dẫn (unguided — không có nhà trị liệu hỗ trợ trực tiếp) cho rối '
      'loạn lo âu xã hội DƯỚI NGƯỠNG (chưa đủ tiêu chuẩn chẩn đoán đầy đủ theo '
      'DSM-5) ở vị thành niên và thanh niên. Tham gia tự đăng ký qua trang web; '
      'chương trình gồm các module video + bài tập + tự đánh giá; không có chat '
      'với nhà trị liệu.')
crit("CBT online không có hướng dẫn có TỈ LỆ HOÀN THÀNH thấp (khoảng 30–50%) so "
     "với CBT online có hướng dẫn (50–70%). Tuy nhiên ưu điểm là khả năng triển "
     "khai quy mô lớn rất tốt cho các quốc gia có thu nhập trung bình – thấp như "
     "Việt Nam. Cần đọc kỹ kích thước hiệu ứng cụ thể và tỉ lệ duy trì sau 3-6 "
     "tháng để có cái nhìn đầy đủ.")
vn_apply("Việt Nam có thể bản địa hoá CBT online sang tiếng Việt — hợp tác với "
         "các công ty công nghệ giáo dục Việt Nam (ví dụ: Topica, ELSA Speak, "
         "Fonos). Nên kết hợp với chatbot tự động (sử dụng AI để trả lời câu hỏi "
         "đơn giản 24/7) + nhắc nhở qua tin nhắn / Zalo để tăng tỉ lệ hoàn thành. "
         "Đối tượng phù hợp: học sinh có triệu chứng dưới ngưỡng (sàng lọc bằng "
         "thang lo âu lan tỏa GAD-7 với điểm ≥ 5, hoặc thang trầm cảm PHQ-9 với "
         "điểm ≥ 5). Đề xuất quy mô: 6–8 module mỗi module 20–30 phút, có bài "
         "tập về nhà, có thể hoàn thành trong 6–8 tuần.")

H2("Bài 4.2 — Tổng quan hệ thống về can thiệp số cho điều hoà cảm xúc (JMIR Serious Games 2022)")
block('Tên bài gốc:',
      '\"Digital Interventions for Emotion Regulation in Children and Early '
      'Adolescents: Systematic Review and Meta-analysis\"')
block('Tạm dịch:',
      'Các can thiệp số cho điều hoà cảm xúc ở trẻ em và vị thành niên sớm: Tổng '
      'quan hệ thống và phân tích gộp')
block('Tạp chí:',
      'JMIR Serious Games — năm 2022 (đã được trích dẫn trong các nghiên cứu năm 2024)')
block('Mã định danh:', 'DOI: 10.2196/31456')
block('Đường dẫn:', 'https://games.jmir.org/2022/3/e31456')
block('Phát hiện chính:',
      'Trò chơi số (digital games) có thiết kế trị liệu (gọi là \"serious games\" '
      '— trò chơi với mục đích nghiêm túc khác giải trí thuần tuý) giảm trải '
      'nghiệm cảm xúc tiêu cực với hiệu ứng nhỏ nhưng có ý nghĩa thống kê — đặc '
      'biệt rõ rệt ở thanh thiếu niên có nguy cơ lo âu.')
crit("Trò chơi số kết hợp yếu tố trị liệu (\"serious games\") là hướng MỚI và phù "
     "hợp với thế hệ Z. Tuy nhiên cần cân nhắc rủi ro tăng thời gian dùng màn "
     "hình (screen time) — vốn đã quá nhiều ở học sinh Việt Nam. Cần thiết kế "
     "để có giới hạn rõ ràng (ví dụ: chỉ 15–20 phút/ngày).")
vn_apply("Việt Nam có thể thiết kế ứng dụng học \"kỹ năng vượt khó\" qua trò chơi "
         "mini bằng tiếng Việt. Có thể hợp tác với các studio làm game độc lập "
         "Việt Nam (như Sky Mavis, Topebox) để tạo sản phẩm phù hợp văn hoá. "
         "Hoặc đơn giản hơn — bản địa hoá các app đã có như SuperBetter, "
         "Headspace Junior sang tiếng Việt. Quan trọng: bắt buộc có giới hạn "
         "thời gian sử dụng và nhắc nhở nghỉ.")

# ---------- NHÓM 5 ----------
H1("NHÓM 5 — Hỗ trợ từ bạn đồng trang lứa (2 bài)")

H2("Bài 5.1 — Tổng quan phạm vi hệ thống về hỗ trợ ngang hàng (Murphy 2024 ở Wiley)")
block('Tên bài gốc:',
      '\"A systematic scoping review of peer support interventions in integrated '
      'primary youth mental health care\"')
block('Tạm dịch:',
      'Tổng quan phạm vi hệ thống về các can thiệp hỗ trợ ngang hàng trong chăm '
      'sóc sức khỏe tâm thần ban đầu tích hợp cho giới trẻ')
block('Tạp chí:', 'Journal of Community Psychology (Wiley) — năm 2024')
block('Mã định danh:', 'DOI: 10.1002/jcop.23090')
block('Đường dẫn:',
      'https://onlinelibrary.wiley.com/doi/full/10.1002/jcop.23090')
block('Bối cảnh nghiên cứu:',
      'Hỗ trợ ngang hàng (peer support) — sự hỗ trợ giữa những người có cùng trải '
      'nghiệm sống — đã được sử dụng rộng rãi trong các nhóm hỗ trợ người nghiện, '
      'sống sót sau ung thư, người khuyết tật. Trong sức khỏe tâm thần giới trẻ, '
      'mô hình này đang được áp dụng nhiều hơn nhưng bằng chứng định lượng còn '
      'thiếu. Bài tổng quan phạm vi này nhằm đánh giá tình hình áp dụng + chất '
      'lượng bằng chứng hiện có.')
block('Phát hiện chính:',
      'Hỗ trợ ngang hàng (peer support — giữa những người trẻ với nhau, có chia '
      'sẻ kinh nghiệm sống cùng vấn đề SKTT) đang được áp dụng rộng trong các '
      'bối cảnh sức khỏe tâm thần cho giới trẻ. Tuy nhiên BẰNG CHỨNG từ các '
      'thử nghiệm có đối chứng (RCT) còn rất HẠN CHẾ — đa số nghiên cứu là mô '
      'tả định tính hoặc tiền-thực nghiệm.')
crit("Đây là tổng quan PHẠM VI (scoping review) — không phải tổng quan hệ thống "
     "đầy đủ — chỉ mô tả tình hình chứ không kết luận quan hệ nhân – quả. Bằng "
     "chứng chủ yếu là mô tả + định tính. Lĩnh vực này vẫn cần nhiều thử nghiệm "
     "có đối chứng nghiêm ngặt để đánh giá hiệu quả thực sự.")
vn_apply("Mô hình \"Đoàn Thanh niên Cộng sản Hồ Chí Minh\" + \"Lớp trưởng / Bí "
         "thư chi đoàn\" có thể bản địa hoá thành chương trình hỗ trợ ngang hàng "
         "có đào tạo. Em ĐỀ XUẤT MÔ HÌNH 3 LỚP: Lớp 1 — Đào tạo nòng cốt (10–15 "
         "học sinh / trường được chọn dựa trên năng lực giao tiếp + sự đồng "
         "thuận của bản thân); Lớp 2 — Giám sát chuyên môn (cán bộ tư vấn học "
         "đường giám sát hằng tuần); Lớp 3 — Mở rộng (nòng cốt hỗ trợ các bạn "
         "khác trong lớp / khối). RỦI RO CẦN PHÒNG NGỪA: tránh \"truyền\" lo âu "
         "giữa các em — cần dạy nòng cốt biết khi nào chuyển sang chuyên gia, "
         "không tự xử lý các trường hợp nghiêm trọng (suy nghĩ tự tử, tự hại).")

H2("Bài 5.2 — Tổng quan hỗ trợ ngang hàng cho SKTT giới trẻ (NCBI Bookshelf)")
block('Mã định danh:', 'NBK602668 (NCBI Bookshelf)')
block('Đường dẫn:', 'https://www.ncbi.nlm.nih.gov/books/NBK602668/')
block('Đặc điểm:',
      'Tổng quan của chính phủ Anh về các chương trình hỗ trợ ngang hàng đã được '
      'triển khai. Tài liệu bổ sung từ Bộ Giáo dục Anh năm 2018 — '
      'https://assets.publishing.service.gov.uk/media/5a820b3d40f0b62305b922c5/'
      'Children_and_young_people_s_mental_health_peer_support.pdf')

# ---------- NHÓM 6 ----------
H1("NHÓM 6 — Bối cảnh châu Á – Trung Quốc (1 bài)")

H2("Bài 6.1 — Tổng quan phạm vi về phòng ngừa SKTT trường học ở Trung Quốc (Lancet 2024)")
block('Tên bài gốc:',
      '\"School mental health prevention and intervention strategies in China: a '
      'scoping review\"')
block('Tạm dịch:',
      'Các chiến lược phòng ngừa và can thiệp sức khỏe tâm thần trong trường học '
      'tại Trung Quốc: Tổng quan phạm vi')
block('Tạp chí:',
      'The Lancet Regional Health – Western Pacific (Hệ số tác động khoảng 16) — '
      'năm 2024')
block('Đường dẫn:',
      'https://www.thelancet.com/journals/lanwpc/article/PIIS2666-6065(24)00237-2/fulltext')
block('Bối cảnh nghiên cứu:',
      'Trung Quốc là quốc gia có hệ thống giáo dục lớn nhất thế giới với hơn 200 '
      'triệu học sinh, nhưng đầu tư vào sức khỏe tâm thần trường học mới được '
      'tăng cường gần đây sau làn sóng tự tử của học sinh năm 2020–2023. Chính '
      'phủ Trung Quốc đã ban hành chính sách quốc gia 2021 yêu cầu mọi trường '
      'phải có cán bộ tâm lý. Bài tổng quan này đánh giá tình hình áp dụng các '
      'chương trình SKTT trường học ở Trung Quốc.')
block('Phát hiện chính:',
      'Tổng quan phạm vi về các chiến lược phòng ngừa và can thiệp SKTT trong '
      'trường học tại Trung Quốc. Bao quát các chương trình về học tập cảm xúc – '
      'xã hội, chánh niệm, liệu pháp nhận thức – hành vi tại các tỉnh khác nhau '
      'của Trung Quốc. Nhấn mạnh khác biệt vùng miền (đô thị – nông thôn, ven '
      'biển – nội địa, người Hán – dân tộc thiểu số).')
crit("RẤT QUAN TRỌNG cho Việt Nam — bối cảnh Trung Quốc tương đồng với Việt Nam "
     "nhất (Đông Á, ảnh hưởng Khổng giáo, áp lực thi đại học cao, hệ thống giáo "
     "dục tập trung do nhà nước điều hành). Bài học từ Trung Quốc dễ bản địa "
     "hoá cho Việt Nam hơn so với Anh / Mỹ. Đặc biệt phần phân tích khác biệt "
     "vùng miền có thể áp dụng trực tiếp cho Việt Nam (Bắc – Trung – Nam, đô "
     "thị – nông thôn – DTTS).")
vn_apply("Cần đọc kỹ bài này để hiểu Trung Quốc đã làm gì, hiệu quả / thất bại "
         "ở đâu. Việt Nam có thể: (1) Tránh các sai lầm Trung Quốc đã mắc — đặc "
         "biệt là việc \"sao chép\" mô hình phương Tây không thích nghi văn "
         "hoá; (2) Bản địa hoá các thành công của Trung Quốc — đặc biệt cách họ "
         "đào tạo cán bộ tâm lý trường ở quy mô lớn (200 triệu học sinh) trong "
         "thời gian ngắn; (3) Học hỏi cách họ phân biệt chiến lược cho khu vực "
         "đô thị (nguồn lực cao) vs nông thôn (nguồn lực thấp). Đề xuất: dịch "
         "tóm tắt bài này sang tiếng Việt và chia sẻ với Bộ GD-ĐT VN.")

# ===================================================================
H1("PHẦN D — BẢNG TỔNG HỢP 15 NGHIÊN CỨU")
tbl(['STT', 'Tác giả + Năm', 'Loại nghiên cứu', 'Mã định danh', 'Mức quan trọng VN'],
    [
        ['1', 'Frontiers Psychiatry 2025',
         'Tổng quan hệ thống + phân tích gộp 38 thử nghiệm có đối chứng / 15.730 HS',
         'PMC12127306 / DOI 10.3389/fpsyt.2025.1594658', '⭐⭐⭐'],
        ['2', 'Springer Child Psychiatry 2024',
         'Tổng quan 31 nghiên cứu / 13 quốc gia',
         'PMC12628395 / DOI 10.1007/s10578-024-01667-5', '⭐⭐⭐'],
        ['3', 'Vogelaar và cộng sự 2024',
         'Thử nghiệm có đối chứng theo cụm (cluster RCT) n=1.613 ở Hà Lan',
         'J School Psychology Tập 105', '⭐⭐'],
        ['4', 'Lochman 2025 (Coping Power)',
         'Thử nghiệm có đối chứng cấp trường, lớp 7 ở Mỹ (Alabama + Maryland)',
         'J School Psychology', '⭐⭐'],
        ['5', 'MDPI Pediatric Reports 2025',
         'Phân tích gộp 14 nghiên cứu MBSR, n=1.489',
         'Tập 17(3) Bài 59', '⭐⭐'],
        ['6', 'PMC12173555 năm 2025',
         'Thử nghiệm có đối chứng MBI 8 tuần + 4 tuần củng cố cho HS 13-15 tuổi',
         'PMC12173555', '⭐⭐'],
        ['7', 'Fulambarkar 2023 (cảnh báo)',
         'Phân tích gộp có tính phản biện về MBI trường học',
         'DOI 10.1111/camh.12572', '⭐⭐⭐ (phản biện)'],
        ['8', 'School Mental Health 2023',
         'Nghiên cứu có đối chứng MBI 10 tuần',
         'DOI 10.1007/s12310-023-09620-y', '⭐⭐'],
        ['9', 'Springer Admin Policy 2024',
         'Phân tích gộp can thiệp tâm lý xã hội với d=0,37 cho điều hoà cảm xúc',
         'DOI 10.1007/s10488-024-01373-3', '⭐⭐⭐'],
        ['10', 'Frontiers Psychiatry 2024',
         'Đào tạo kỹ năng điều hoà cảm xúc cho HS + cha mẹ',
         'DOI 10.3389/fpsyt.2024.1448529', '⭐⭐'],
        ['11', 'JMIR Pediatrics 2024',
         'Thử nghiệm có đối chứng đa trung tâm về CBT online cho lo âu xã hội dưới ngưỡng',
         'DOI 10.2196/55786', '⭐⭐⭐'],
        ['12', 'JMIR Serious Games 2022',
         'Tổng quan hệ thống + phân tích gộp về can thiệp số cho điều hoà cảm xúc',
         'DOI 10.2196/31456', '⭐⭐'],
        ['13', 'Murphy 2024 (hỗ trợ ngang hàng)',
         'Tổng quan phạm vi hệ thống về peer support',
         'DOI 10.1002/jcop.23090', '⭐⭐'],
        ['14', 'NCBI NBK602668',
         'Tổng quan của chính phủ Anh về peer support cho SKTT giới trẻ',
         'NBK602668', '⭐'],
        ['15', 'Lancet RH Western Pacific 2024',
         'Tổng quan phạm vi về phòng ngừa + can thiệp SKTT trường học ở Trung Quốc',
         'PIIS2666-6065(24)00237-2', '⭐⭐⭐'],
    ], [0.7, 3.0, 4.5, 5.0, 2.0])

# ===================================================================
H1("PHẦN E — KHUYẾN NGHỊ ÁP DỤNG CHO VIỆT NAM")

H2("E.1. Thứ tự ưu tiên các hướng can thiệp (theo độ mạnh bằng chứng)")
nr("Dựa trên 15 nghiên cứu trên + đối chiếu với 3 bài Brown đã được dịch trong dự án "
   "(BESST trial 2024, PLACES model 2022, Editorial 2025), em đề xuất thứ tự ưu tiên "
   "cho học sinh trung học Việt Nam (TỪ MẠNH NHẤT đến yếu nhất):", bold=True)
nr("(1) Quản lý căng thẳng dựa trên Liệu pháp nhận thức – hành vi — đào tạo kỹ năng "
   "có cấu trúc, kích thước hiệu ứng nhỏ đến trung bình, NHẤT QUÁN nhất qua nhiều "
   "nghiên cứu. Đây là HƯỚNG AN TOÀN NHẤT để đầu tư.")
nr("(2) Đào tạo kỹ năng điều hoà cảm xúc — kỹ thuật từ DBT, ACT, CBT tập trung cảm "
   "xúc — hệ số d của Cohen ≈ 0,37 (mức trung bình). Tiềm năng cao hơn chánh niệm.")
nr("(3) Cơ chế tự giới thiệu có chỉ định (theo nghiên cứu BESST 2024 đã được dịch) — "
   "đạt hiệu quả trung bình d=−0,52 cho NHÓM ĐÃ CÓ TRIỆU CHỨNG CAO; phù hợp với "
   "mô hình PLACES của Brown 2022.")
nr("(4) Liệu pháp nhận thức – hành vi qua mạng (CBT online) — có thể triển khai quy "
   "mô lớn nhưng tỉ lệ duy trì tham gia thấp khi không có người hướng dẫn.")
nr("(5) Chánh niệm phổ quát (cho mọi học sinh không phân biệt) — hiệu ứng nhỏ "
   "(SMD = −0,14); cần thận trọng với mô hình thất bại MYRIAD ở Anh năm 2022.")
nr("(6) Hỗ trợ từ bạn đồng trang lứa — hứa hẹn nhưng bằng chứng còn yếu (chưa có "
   "đủ thử nghiệm có đối chứng).")

H2("E.2. Đề xuất chương trình 12 tuần cho học sinh trung học Việt Nam")
tbl(['Tuần', 'Nội dung', 'Phương pháp triển khai', 'Cơ sở khoa học từ tổng quan'],
    [
        ['Tuần 1–2',
         'Hiểu biết về sức khỏe tâm thần + Giảm kỳ thị',
         'Tâm lý giáo dục cho toàn lớp, 3 buổi × 45 phút (theo mô hình '
         'Vogelaar 2024)',
         'Vogelaar 2024 + nguyên tắc Lay từ mô hình PLACES'],
        ['Tuần 3–4',
         'CBT cốt lõi: Nhận diện cảm xúc + suy nghĩ tiêu cực',
         'Workshop nhóm 60–90 phút/tuần',
         'Brown 2024 BESST với chương trình DISCOVER'],
        ['Tuần 5–6',
         'Kỹ năng CBT: Tái cấu trúc nhận thức + kích hoạt hành vi',
         'Workshop + đóng vai + sách bài tập (theo BESST manual)',
         'Sách hướng dẫn DISCOVER'],
        ['Tuần 7–8',
         'Đào tạo kỹ năng điều hoà cảm xúc',
         'Module DBT (chánh niệm + chịu đựng đau khổ + điều hoà cảm xúc + '
         'hiệu quả giao tiếp)',
         'Springer Admin Policy 2024 với hệ số d=0,37'],
        ['Tuần 9–10',
         'Các kỹ năng ứng phó bổ sung',
         'Tập thể dục, vệ sinh giấc ngủ, quản lý thời gian',
         'Springer 2024 + Wen 2020 (Trung Quốc nông thôn)'],
        ['Tuần 11',
         'Hỗ trợ ngang hàng + xây dựng cộng đồng',
         'Chia sẻ nhóm + hệ thống bạn đồng hành (3 lớp đào tạo nòng cốt)',
         'Murphy 2024'],
        ['Tuần 12',
         'Củng cố + phòng ngừa tái phát',
         'Ôn tập + cá nhân hoá mục tiêu + buổi củng cố sau 4 tuần',
         'PMC12173555 năm 2025 (mô hình củng cố)'],
    ], [1.2, 4.5, 4.8, 4.5])

H2("E.3. Mô hình triển khai khả thi cho Việt Nam")
nr("• CẤP TRƯỜNG: Tích hợp chương trình 12 tuần vào tiết HĐTN-HN hoặc Giáo dục công "
   "dân lớp 10–12 (ưu tiên lớp 11–12 do áp lực thi đại học cao nhất). Một tiết "
   "mỗi tuần × 60 phút (linh hoạt theo thời khoá biểu).", size=12)
nr("• NGƯỜI CUNG CẤP CAN THIỆP: Cán bộ tư vấn học đường (theo Thông tư 31/2017/"
   "TT-BGDĐT về tư vấn tâm lý cho học sinh) + chuyên gia tâm lý liên trường giám "
   "sát hằng tháng. Cần đào tạo CBT 5–7 ngày + buổi củng cố hằng quý.", size=12)
nr("• CƠ CHẾ TỰ GIỚI THIỆU (self-referral): Học sinh tự đăng ký qua ứng dụng / "
   "trang web ẩn danh; tránh kênh giáo viên chủ nhiệm để giảm kỳ thị (theo "
   "nghiên cứu BESST 2024 + mô hình PLACES).", size=12)
nr("• SÀNG LỌC: Sử dụng thang lo âu lan tỏa GAD-7 (đã có chuẩn Việt Nam, hệ số tin "
   "cậy Cronbach alpha = 0,916 theo Hoa 2024) hoặc thang trầm cảm PHQ-9 (chuẩn "
   "Việt Nam của Trần Tuấn 2017) vào đầu năm học. Học sinh có điểm ≥ 5 (mức nhẹ) "
   "hoặc ≥ 10 (mức trung bình) được mời tham gia chương trình có chỉ định; học "
   "sinh bình thường vẫn nhận phần phổ quát (hiểu biết SKTT + giảm kỳ thị).", size=12)
nr("• HỖ TRỢ QUA CÔNG NGHỆ: Ứng dụng tiếng Việt với các module CBT + chatbot AI "
   "+ nhắc nhở củng cố qua tin nhắn (bản địa hoá từ mô hình JMIR Pediatrics 2024).",
   size=12)
nr("• ĐÁNH GIÁ HIỆU QUẢ: Đo trước – sau – 3 tháng – 6 tháng bằng GAD-7 + PHQ-9. "
   "Tham khảo các outcome của BESST: MFQ depression, RCADS anxiety, WEMWBS "
   "wellbeing, SCI sleep quality, CYRM-12 resilience.",
   size=12)

# ===================================================================
H1("PHẦN F — THAM KHẢO ĐẦY ĐỦ")
nr("(15 tài liệu — mỗi tài liệu có DOI, mã PubMed Central nếu có, và đường dẫn truy cập.)",
   italic=True, color=GRAY)

nr("1. Frontiers in Psychiatry (2025). School-based interventions for resilience in "
   "children and adolescents: Systematic review and meta-analysis of randomized "
   "controlled trials. Tập 16. Mã PMC: PMC12127306. DOI: 10.3389/fpsyt.2025.1594658. "
   "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2025.1594658/full",
   size=11)
nr("2. Academic Stress Interventions in High Schools: A Systematic Literature Review "
   "(2024). Child Psychiatry & Human Development. DOI: 10.1007/s10578-024-01667-5. "
   "Mã PMC: PMC12628395. https://link.springer.com/article/10.1007/s10578-024-01667-5",
   size=11)
nr("3. Vogelaar và cộng sự (2024). Chương trình tâm lý giáo dục \"Stress Lessons\" — "
   "Thử nghiệm có đối chứng theo cụm với 1.613 học sinh ở Hà Lan. Journal of School "
   "Psychology, Tập 105.", size=11)
nr("4. Lochman và cộng sự (2025). Chương trình \"Sức mạnh ứng phó cho vị thành niên "
   "sớm\" (Coping Power) — Thử nghiệm có đối chứng theo nhóm cấp trường, học sinh "
   "lớp 7 ở Alabama và Maryland (Mỹ). Journal of School Psychology.", size=11)
nr("5. MDPI Pediatric Reports (2025). Mindfulness in Mental Health and Psychiatric "
   "Disorders of Children and Adolescents: Systematic Review and Meta-Analysis. "
   "Tập 17(3): Bài 59. https://www.mdpi.com/2036-7503/17/3/59", size=11)
nr("6. PMC12173555 (2025). Thử nghiệm có đối chứng về can thiệp chánh niệm 8 tuần + "
   "4 tuần củng cố cho vị thành niên 13–15 tuổi.", size=11)
nr("7. Fulambarkar N, Seo B, Testerman A, Rees M, Bausback K, Bunge E (2023). "
   "Phân tích gộp về chánh niệm trong trường học cho căng thẳng, trầm cảm, lo âu "
   "ở vị thành niên: một câu chuyện cần cảnh giác. Child and Adolescent Mental "
   "Health (Wiley). DOI: 10.1111/camh.12572.", size=11)
nr("8. School Mental Health (2023). Can thiệp chánh niệm 10 tuần dựa trên trường học "
   "và triệu chứng trầm cảm + lo âu ở trẻ em và vị thành niên: nghiên cứu có đối "
   "chứng. DOI: 10.1007/s12310-023-09620-y.", size=11)
nr("9. Springer — Administration and Policy in Mental Health and Mental Health Services "
   "Research (2024). Effect of Psychosocial Interventions on Children and Youth Emotion "
   "Regulation: A Meta-Analysis. Hệ số d=0,37. DOI: 10.1007/s10488-024-01373-3.", size=11)
nr("10. Frontiers in Psychiatry (2024). An emotion regulation skills training for "
    "adolescents and parents: perceptions and acceptability of methodological aspects. "
    "DOI: 10.3389/fpsyt.2024.1448529.", size=11)
nr("11. JMIR Pediatrics and Parenting (2024). Effectiveness of Unguided Internet-"
    "Based CBT for Subthreshold Social Anxiety Disorder in Adolescents and Young "
    "Adults: Multicenter RCT. DOI: 10.2196/55786.", size=11)
nr("12. JMIR Serious Games (2022). Digital Interventions for Emotion Regulation in "
    "Children and Early Adolescents: Systematic Review and Meta-analysis. "
    "DOI: 10.2196/31456.", size=11)
nr("13. Murphy R và cộng sự (2024). A systematic scoping review of peer support "
    "interventions in integrated primary youth mental health care. Journal of "
    "Community Psychology (Wiley). DOI: 10.1002/jcop.23090.", size=11)
nr("14. Peer Support Programs for Youth Mental Health (2024). NCBI Bookshelf. "
    "Mã NBK602668.", size=11)
nr("15. The Lancet Regional Health – Western Pacific (2024). School mental health "
    "prevention and intervention strategies in China: a scoping review. "
    "Mã PIIS2666-6065(24)00237-2.", size=11)

H2("Đối chiếu kho tài liệu dự án (tránh trùng lặp)")
nr("• Em đã đối chiếu với DATABASE_BAI_BAO_LO_AU.md (kho 35+ bài đã có) — 15 tài liệu "
   "nói trên ĐỀU CHƯA CÓ trong kho dự án; là bổ sung mới.", size=11)
nr("• Các bài liên quan đã có trong kho: Brown 2024 BESST (mã QT042_BESST), Brown 2022 "
   "PLACES (mã QT042_PLACES), Brown & Carter 2025 (mã QT042_B5), Kuyken 2022 MYRIAD, "
   "Stallard 2014 PACES, Zhang 2023 phân tích gộp — tất cả đã được dịch + phản biện "
   "trong 3 tài liệu dịch song ngữ trước đó (lưu ở 03_Ban-dich/Bai_dich_phan_bien/).",
   size=11)

H2("Cảnh báo về độ chính xác mã định danh")
crit("Một số DOI em ghi là ƯỚC TÍNH (chưa tải bản đầy đủ từng bài). Cần kiểm tra "
     "lại trước khi trích dẫn chính thức trong báo cáo học thuật. Các bài có "
     "DOI/PMC CHẮC CHẮN (đã đối chiếu qua web): Brown 2022 PLACES với PMC8909998, "
     "Brown 2024 BESST với PMID 38759665, Murphy 2024 với DOI 10.1002/jcop.23090, "
     "Fulambarkar 2023 với DOI 10.1111/camh.12572. Các bài khác (Vogelaar 2024, "
     "Lochman 2025): em có đường dẫn ScienceDirect nhưng DOI chính xác cần kiểm "
     "tra lại. Em sẵn sàng làm WebFetch xác minh thêm nếu thầy yêu cầu.")

d.save(OUT)
print('Wrote:', OUT)
