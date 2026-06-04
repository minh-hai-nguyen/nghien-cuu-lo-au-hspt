# -*- coding: utf-8 -*-
"""
kg_crosscheck_all_v1.py — Cross-check KG v2 vs ALL translations + summaries.
Phát hiện mâu thuẫn giữa:
  (a) KG effect sizes vs numbers trong bản dịch/tóm tắt
  (b) Summary vs Translation cho cùng 1 paper
  (c) Numbers xuất hiện trong nhiều files nhưng khác giá trị

Output: JSON report + console summary
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import networkx as nx
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
KG_FILE = os.path.join(os.path.dirname(__file__), 'kg_data', 'kg_v2.graphml')
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
BD_DIR = os.path.join(ROOT, '03_Ban-dich')

# ============================================================
# 1. Load KG v2
# ============================================================
print('Loading KG v2...')
G = nx.read_graphml(KG_FILE)
print(f'  Nodes: {G.number_of_nodes()}, Edges: {G.number_of_edges()}')

# Extract effect sizes from KG
kg_effects = {}  # {paper_id: [(metric, value, context), ...]}
for n, d in G.nodes(data=True):
    ntype = d.get('type', '')
    if ntype == 'EffectSize':
        label = d.get('label', n)
        context = d.get('context', '')
        # Find connected paper
        for pred in G.predecessors(n):
            pd = G.nodes[pred]
            if pd.get('type') == 'Paper':
                pid = pd.get('label', pred)
                if pid not in kg_effects:
                    kg_effects[pid] = []
                kg_effects[pid].append((label, context))
                break
        for succ in G.successors(n):
            sd = G.nodes[succ]
            if sd.get('type') == 'Paper':
                pid = sd.get('label', succ)
                if pid not in kg_effects:
                    kg_effects[pid] = []
                kg_effects[pid].append((label, context))
                break

# Also try edge-based extraction
for u, v, d in G.edges(data=True):
    rel = d.get('relation', d.get('label', ''))
    un = G.nodes[u]
    vn = G.nodes[v]
    if un.get('type') == 'Paper' and vn.get('type') == 'EffectSize':
        pid = un.get('label', u)
        es_label = vn.get('label', v)
        es_ctx = vn.get('context', '')
        if pid not in kg_effects:
            kg_effects[pid] = []
        kg_effects[pid].append((es_label, es_ctx))

print(f'  Papers with effect sizes in KG: {len(kg_effects)}')
total_es = sum(len(v) for v in kg_effects.values())
print(f'  Total effect size entries: {total_es}')

# ============================================================
# 2. Load all summaries + translations
# ============================================================
def read_docx(path):
    """Extract all text from docx."""
    try:
        d = Document(path)
        lines = []
        for p in d.paragraphs:
            if p.text.strip():
                lines.append(p.text)
        for t in d.tables:
            for r in t.rows:
                row = ' | '.join(c.text.strip() for c in r.cells)
                if row.strip():
                    lines.append(row)
        return '\n'.join(lines)
    except Exception as e:
        return f'ERROR: {e}'

def get_canonical_id(filename):
    """Extract VN001 or QT001 from filename."""
    m = re.match(r'(VN\d{3}|QT\d{3})', filename)
    return m.group(1) if m else None

print('\nLoading summaries...')
summaries = {}
for f in sorted(os.listdir(TT_DIR)):
    if f.endswith('.docx') and not f.startswith('~') and not f.startswith('_'):
        cid = get_canonical_id(f)
        if cid:
            summaries[cid] = read_docx(os.path.join(TT_DIR, f))
print(f'  Loaded: {len(summaries)} summaries')

print('Loading translations...')
translations = {}
for f in sorted(os.listdir(BD_DIR)):
    if f.endswith('.docx') and not f.startswith('~') and not f.startswith('_'):
        cid = get_canonical_id(f)
        if cid:
            translations[cid] = read_docx(os.path.join(BD_DIR, f))
# Also check subdirs
for subdir in os.listdir(BD_DIR):
    subpath = os.path.join(BD_DIR, subdir)
    if os.path.isdir(subpath):
        for f in os.listdir(subpath):
            if f.endswith('.docx') and not f.startswith('~'):
                cid = get_canonical_id(f)
                if cid and cid not in translations:
                    translations[cid] = read_docx(os.path.join(subpath, f))
            elif f.endswith('.md') and not f.startswith('~'):
                cid = get_canonical_id(f)
                if cid and cid not in translations:
                    try:
                        with open(os.path.join(subpath, f), encoding='utf-8') as fh:
                            translations[cid] = fh.read()
                    except:
                        pass
print(f'  Loaded: {len(translations)} translations')

# ============================================================
# 3. Extract numbers from text
# ============================================================
NUM_PATTERNS = [
    (r'OR\s*=?\s*(\d+[,.]?\d*)', 'OR'),
    (r'aOR\s*=?\s*(\d+[,.]?\d*)', 'aOR'),
    (r'AOR\s*=?\s*(\d+[,.]?\d*)', 'AOR'),
    (r'β\s*=?\s*[-−]?(\d+[,.]?\d*)', 'beta'),
    (r'[Cc]ohen\s*d\s*=?\s*[-−]?(\d+[,.]?\d*)', 'Cohen_d'),
    (r'\bd\s*=\s*[-−]?(\d+[,.]?\d*)', 'd'),
    (r'[Hh]edges?\s*g\s*=?\s*[-−]?(\d+[,.]?\d*)', 'Hedges_g'),
    (r'\bg\s*=\s*[-−]?(\d+[,.]?\d*)', 'g'),
    (r'SMD\s*=?\s*[-−]?(\d+[,.]?\d*)', 'SMD'),
    (r'SUCRA\s*=?\s*(\d+[,.]?\d*)', 'SUCRA'),
    (r'n\s*=\s*([\d.,]+)', 'n'),
]

def extract_numbers(text):
    out = []
    for pat, metric in NUM_PATTERNS:
        for m in re.finditer(pat, text):
            val = m.group(1).replace(',', '.')
            try:
                fval = float(val.replace('.', '', val.count('.')-1) if val.count('.') > 1 else val)
            except:
                fval = None
            ctx_start = max(0, m.start() - 30)
            ctx_end = min(len(text), m.end() + 30)
            ctx = text[ctx_start:ctx_end].replace('\n', ' ').strip()
            out.append({'metric': metric, 'value': val, 'fval': fval, 'context': ctx})
    return out

# ============================================================
# 4. Cross-check: Summary vs Translation
# ============================================================
print('\n=== CROSS-CHECK: Summary vs Translation ===')
contradictions = []

all_ids = sorted(set(summaries.keys()) | set(translations.keys()))
print(f'Total papers with at least 1 file: {len(all_ids)}')

for cid in all_ids:
    sum_text = summaries.get(cid, '')
    trans_text = translations.get(cid, '')

    if not sum_text or not trans_text:
        continue  # can't compare if one is missing

    sum_nums = extract_numbers(sum_text)
    trans_nums = extract_numbers(trans_text)

    # For each number in summary, check if it appears in translation
    for sn in sum_nums:
        if sn['fval'] is None or sn['fval'] == 0:
            continue
        # Look for same metric + similar value in translation
        found_match = False
        for tn in trans_nums:
            if tn['fval'] is None:
                continue
            if sn['metric'] == tn['metric']:
                # Allow 1% tolerance for float comparison
                if abs(sn['fval'] - tn['fval']) < 0.015:
                    found_match = True
                    break
            # Also check cross-metric (aOR vs AOR vs OR)
            if sn['metric'].lower().replace('a','').replace('o','').replace('r','') == \
               tn['metric'].lower().replace('a','').replace('o','').replace('r',''):
                if abs(sn['fval'] - tn['fval']) < 0.015:
                    found_match = True
                    break

        if not found_match and sn['metric'] in ('OR', 'aOR', 'AOR', 'beta', 'Cohen_d', 'd', 'Hedges_g', 'g', 'SMD', 'SUCRA'):
            # Check if the raw value string appears anywhere in translation
            val_str = sn['value']
            val_comma = val_str.replace('.', ',')
            val_dot = val_str.replace(',', '.')
            if val_str in trans_text or val_comma in trans_text or val_dot in trans_text:
                found_match = True

            if not found_match:
                contradictions.append({
                    'type': 'summary_not_in_translation',
                    'paper': cid,
                    'metric': sn['metric'],
                    'value': sn['value'],
                    'sum_context': sn['context'],
                    'note': f'{sn["metric"]}={sn["value"]} in summary but not found in translation'
                })

print(f'  Summary-only numbers (not in translation): {len(contradictions)}')

# ============================================================
# 5. Cross-check: KG effect sizes vs source files
# ============================================================
print('\n=== CROSS-CHECK: KG vs Source Files ===')
kg_mismatches = []

for pid, effects in kg_effects.items():
    source_text = translations.get(pid, '') + '\n' + summaries.get(pid, '')
    if not source_text.strip():
        kg_mismatches.append({
            'type': 'kg_no_source',
            'paper': pid,
            'note': f'KG has {len(effects)} effect sizes but no source files found'
        })
        continue

    for es_label, es_ctx in effects:
        # Extract numeric value from KG label
        nums = re.findall(r'[\d]+[.,][\d]+', es_label)
        if not nums:
            continue
        for num in nums:
            num_dot = num.replace(',', '.')
            num_comma = num.replace('.', ',')
            if num not in source_text and num_dot not in source_text and num_comma not in source_text:
                kg_mismatches.append({
                    'type': 'kg_not_in_source',
                    'paper': pid,
                    'kg_label': es_label,
                    'kg_context': es_ctx[:100],
                    'searched_value': num,
                    'note': f'KG effect size {num} not found in summary/translation'
                })

print(f'  KG values not found in source: {len(kg_mismatches)}')

# ============================================================
# 6. Cross-check: Same fact different values across papers
# ============================================================
print('\n=== CROSS-CHECK: Cross-paper number references ===')
cross_issues = []

# Build index: which papers mention which other papers' IDs
for cid in all_ids:
    text = (summaries.get(cid, '') + '\n' + translations.get(cid, '')).strip()
    if not text:
        continue
    # Find references to other papers
    refs = re.findall(r'(VN\d{3}|QT\d{3})', text)
    refs = [r for r in set(refs) if r != cid]  # exclude self

    for ref_id in refs:
        # Get numbers mentioned in context of ref_id in this paper
        # Find paragraphs mentioning ref_id
        for line in text.split('\n'):
            if ref_id in line:
                nums_here = extract_numbers(line)
                if not nums_here:
                    continue
                # Check if these numbers match what's in the ref_id's own files
                ref_text = (summaries.get(ref_id, '') + '\n' + translations.get(ref_id, '')).strip()
                if not ref_text:
                    continue
                for nh in nums_here:
                    if nh['fval'] is None or nh['metric'] == 'n':
                        continue
                    val_str = nh['value']
                    val_comma = val_str.replace('.', ',')
                    val_dot = val_str.replace(',', '.')
                    if val_str not in ref_text and val_comma not in ref_text and val_dot not in ref_text:
                        cross_issues.append({
                            'type': 'cross_paper_mismatch',
                            'citing_paper': cid,
                            'referenced_paper': ref_id,
                            'metric': nh['metric'],
                            'value': nh['value'],
                            'context': nh['context'],
                            'note': f'{cid} cites {ref_id} with {nh["metric"]}={nh["value"]} but value not found in {ref_id} source'
                        })

print(f'  Cross-paper mismatches: {len(cross_issues)}')

# ============================================================
# 7. Report
# ============================================================
all_issues = contradictions + kg_mismatches + cross_issues
print(f'\n{"="*70}')
print(f'TOTAL ISSUES: {len(all_issues)}')
print(f'  Summary vs Translation: {len(contradictions)}')
print(f'  KG vs Source: {len(kg_mismatches)}')
print(f'  Cross-paper: {len(cross_issues)}')
print(f'{"="*70}')

# Print top issues
if contradictions:
    print(f'\n--- Top Summary vs Translation issues (first 15) ---')
    for c in contradictions[:15]:
        print(f'  [{c["paper"]}] {c["metric"]}={c["value"]} — {c["note"]}')
        print(f'    Context: {c["sum_context"][:120]}')

if kg_mismatches:
    print(f'\n--- KG vs Source issues (first 15) ---')
    for c in kg_mismatches[:15]:
        if c['type'] == 'kg_no_source':
            print(f'  [{c["paper"]}] {c["note"]}')
        else:
            print(f'  [{c["paper"]}] KG: {c["kg_label"][:60]} — searched: {c["searched_value"]}')

if cross_issues:
    print(f'\n--- Cross-paper issues (first 15) ---')
    for c in cross_issues[:15]:
        print(f'  [{c["citing_paper"]}→{c["referenced_paper"]}] {c["metric"]}={c["value"]}')
        print(f'    Context: {c["context"][:120]}')

# Save JSON
out_path = os.path.join(os.path.dirname(__file__), 'kg_crosscheck_all_report.json')
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump({
        'total_papers': len(all_ids),
        'summaries_loaded': len(summaries),
        'translations_loaded': len(translations),
        'kg_papers_with_es': len(kg_effects),
        'total_issues': len(all_issues),
        'summary_vs_translation': len(contradictions),
        'kg_vs_source': len(kg_mismatches),
        'cross_paper': len(cross_issues),
        'issues': all_issues,
    }, f, ensure_ascii=False, indent=2)

print(f'\nFull report saved: {out_path}')
print('Done.')
