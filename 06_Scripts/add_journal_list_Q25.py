# -*- coding: utf-8 -*-
"""
Thêm danh sách 12 tạp chí Q2-Q3 có thể gửi cho paper Q2.5 này vào cuối VN + EN.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
VN = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v3.docx"
EN = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_EN_v3.docx"


def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts"); rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_para(doc, text, size=12, bold=False, italic=False, align=None, indent=None,
             space_before=3, space_after=3):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before); pf.space_after = Pt(space_after)
    pf.line_spacing = 1.15
    if indent is not None: pf.first_line_indent = Cm(indent)
    if align == "center": p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify": p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_format(r, size=size, bold=bold, italic=italic)
    return p


def add_h1(doc, text):
    return add_para(doc, text, size=13, bold=True, space_before=14, space_after=6)


def add_h2(doc, text):
    return add_para(doc, text, size=12, bold=True, italic=True, space_before=8, space_after=4)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Light Grid"
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]; cell.text = ""
        r = cell.paragraphs[0].add_run(h); set_run_format(r, size=10, bold=True)
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]; cell.text = ""
            r = cell.paragraphs[0].add_run(str(v)); set_run_format(r, size=10)
    return t


# ============================================================
# 12 tạp chí — sắp xếp theo chiến lược nộp (Q2 sweet spot → Q3 backup)
# ============================================================

JOURNALS = [
    # (Rank, Name, ISSN, Publisher, IF, Quartile, APC, Turnaround, Scope-fit, Vietnamese-authors, Notes)
    (1, "Frontiers in Psychiatry", "1664-0640", "Frontiers", "3,10",
     "Q1 SJR / Q2 WoS", "~3.150 CHF (waiver LMIC)", "~77 ngày", "⭐⭐⭐⭐⭐",
     "Có", "Lựa chọn ưu tiên #1 — tốc độ + uy tín"),
    (2, "Asian Journal of Psychiatry", "1876-2018", "Elsevier", "2,77",
     "Q1 SJR", "Tùy editor (waiver LMIC)", "Không rõ", "⭐⭐⭐⭐⭐",
     "Mạnh", "Cultural fit cao nhất — bài hợp khu vực Á"),
    (3, "Child Psychiatry & Human Development", "0009-398X", "Springer", "2,20",
     "Q2 SJR", "~$2-3k", "Không rõ", "⭐⭐⭐⭐⭐",
     "Vừa", "Springer + PubMed indexed; chấp nhận SEM"),
    (4, "BMC Psychology", "2050-7283", "Springer Nature (BMC)", "3,40",
     "Q1 SJR", "€1.390 (OA)", "Không rõ", "⭐⭐⭐⭐",
     "Có", "OA rẻ nhất nhóm Q1-Q2; PMC indexed"),
    (5, "Children and Youth Services Review", "0190-7409", "Elsevier", "2,14",
     "Q1 (policy lane)", "Subscription (miễn phí cho tác giả)", "~8,7 tuần", "⭐⭐⭐⭐",
     "Vừa", "Tốt cho hàm ý chính sách + dịch vụ học đường"),
    (6, "Frontiers in Psychology", "1664-1078", "Frontiers", "3,27",
     "Q2", "~$2.500", "~77 ngày", "⭐⭐⭐⭐",
     "Có", "Phù hợp khi nhấn psychology developmental"),
    (7, "Journal of Mental Health", "0963-8237", "Taylor & Francis", "~2,5",
     "Q2", "Hybrid (tùy chọn OA)", "Không rõ", "⭐⭐⭐",
     "Ít", "T&F hybrid model — đọc kỹ chính sách trước nộp"),
    (8, "Journal of School Health", "0022-4391", "Wiley", "2,04",
     "Q2 SJR", "Optional OA fee", "Không rõ", "⭐⭐⭐",
     "Ít", "Nhấn implications học đường"),
    (9, "School Mental Health", "1866-2625", "Springer", "~3,0",
     "Q2", "~$3k", "Không rõ", "⭐⭐⭐⭐",
     "Vừa", "Đúng niche school-based MH"),
    (10, "PLOS ONE", "1932-6203", "PLOS", "2,60",
     "Q2 WoS / Q1 PubMed", "$1.695 (OA)", "~185 ngày", "⭐⭐⭐",
     "Vừa", "Acceptance rate cao ~50%; broad scope"),
    (11, "Cambridge Prisms: Global Mental Health", "2054-4251", "Cambridge UP", "~3,5",
     "Q2", "$2.500 (OA)", "Không rõ", "⭐⭐⭐⭐",
     "Có", "LMIC focus — Vietnamese authors welcome"),
    (12, "Vulnerable Children and Youth Studies", "1745-0128", "Taylor & Francis", "~1,3",
     "Q3", "Hybrid", "Không rõ", "⭐⭐⭐",
     "Ít", "Backup khi cascade reject từ tier cao hơn"),
]


# ============================================================
# VN content
# ============================================================
def add_vn_section(doc):
    add_h1(doc, "Phụ lục bổ sung — Danh sách 12 tạp chí mục tiêu (Q2 đến Q3)")

    add_para(doc,
             "Danh sách dưới đây tổng hợp 12 tạp chí quốc tế phù hợp cho bài Q2.5 này, được sắp "
             "xếp theo thứ tự ưu tiên nộp. Mỗi tạp chí đi kèm chỉ số tác động (Impact Factor) "
             "năm 2024, phân khúc Quartile, phí xuất bản open access (APC), thời gian phản biện "
             "ước tính, mức phù hợp scope (đánh giá 1–5 sao) và ghi chú chiến thuật.",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "Bảng — 12 tạp chí mục tiêu")
    add_table(doc,
              ["#", "Tạp chí", "IF 2024", "Quartile", "APC", "Turnaround", "Fit", "Ghi chú"],
              [
                  (str(j[0]), j[1], j[4], j[5], j[6], j[7], j[8], j[10])
                  for j in JOURNALS
              ])

    add_h2(doc, "Chiến thuật nộp (submission cascade)")
    cascade_text = [
        "Tier 1 — uy tín cao, tốc độ nhanh (Tuần 0–12). Nộp Frontiers in Psychiatry trước; "
        "quyết định trong ~77 ngày. Nếu reject, chuyển sang Asian Journal of Psychiatry (lựa "
        "chọn cultural fit tốt nhất).",
        "Tier 2 — Springer/Wiley backup (Tuần 12–24). Nếu cả Tier 1 reject, nộp Child "
        "Psychiatry & Human Development (Springer) hoặc School Mental Health (Springer) — đều "
        "có ecosystem phản biện rõ ràng.",
        "Tier 3 — open access broad (Tuần 24–36). Nếu Tier 2 reject, nộp BMC Psychology "
        "(€1.390 — rẻ nhất) hoặc Frontiers in Psychology. Cambridge Prisms: Global Mental "
        "Health đặc biệt phù hợp do nhấn mạnh LMIC.",
        "Tier 4 — policy/services hoặc Q3 (Tuần 36+). Nếu vẫn reject, chuyển sang Children and "
        "Youth Services Review (policy lane), Journal of School Health (nhấn implications học "
        "đường) hoặc Vulnerable Children and Youth Studies (Q3 backup an toàn).",
    ]
    for c in cascade_text:
        add_para(doc, "• " + c, size=12, indent=0.3, align="justify")

    add_h2(doc, "Cảnh báo cần tránh")
    warnings = [
        "Tránh MDPI International Journal of Environmental Research and Public Health "
        "(IJERPH) — Clarivate đã loại tạp chí này khỏi Web of Science năm 2024 do quan ngại "
        "về citation inflation.",
        "Cẩn thận với các tạp chí có model hybrid (kết hợp subscription + OA) — đọc kỹ chính "
        "sách phí trước khi nộp.",
        "Tránh các tạp chí có acceptance rate < 10% nếu bài chưa qua editing tiếng Anh "
        "chuyên nghiệp.",
    ]
    for w in warnings:
        add_para(doc, "• " + w, size=12, indent=0.3, align="justify")

    add_h2(doc, "Lưu ý chuẩn bị nộp")
    notes = [
        "Tất cả tạp chí Q2-Q3 trên đều yêu cầu cover letter ngắn (≤ 1 trang) nêu rõ "
        "đóng góp mới + lý do phù hợp với scope.",
        "Bài Q2.5 này nên được edit ngôn ngữ bởi native English speaker (Editage, Enago, "
        "American Journal Experts) trước khi nộp Tier 1, vì các tạp chí Frontiers + Springer + "
        "Elsevier có thể desk-reject vì lý do ngôn ngữ.",
        "Khai báo công cụ AI (LLM) trong cover letter và trong mục Phương pháp nghiên cứu, "
        "tuân theo yêu cầu hiện hành của các nhà xuất bản lớn (Elsevier, Springer, Wiley, "
        "Frontiers đều đã có chính sách riêng từ 2024).",
        "Bộ ba bài đồng hành (Validation của Công và Đào 2026 + 2 bài VJES) nên được cite "
        "rõ ràng trong Introduction và Methods để editor không nghi salami slicing.",
    ]
    for n in notes:
        add_para(doc, "• " + n, size=12, indent=0.3, align="justify")


# ============================================================
# EN content
# ============================================================
def add_en_section(doc):
    add_h1(doc, "Supplementary appendix — Twelve target journals (Q2 to Q3)")

    add_para(doc,
             "The table below lists twelve international journals appropriate for this Q2.5 "
             "manuscript, ordered by submission priority. Each row reports the 2024 Impact "
             "Factor, quartile, open-access fee (APC), estimated turnaround, scope-fit rating "
             "(1–5 stars), and a brief strategic note.",
             size=12, align="justify", indent=0.5)

    add_h2(doc, "Table — Twelve target journals")
    add_table(doc,
              ["#", "Journal", "IF 2024", "Quartile", "APC", "Turnaround", "Fit", "Notes"],
              [
                  (str(j[0]), j[1], j[4], j[5], j[6], j[7], j[8], j[10])
                  for j in JOURNALS
              ])

    add_h2(doc, "Submission cascade strategy")
    cascade = [
        "Tier 1 — prestige + speed (weeks 0–12). Submit to Frontiers in Psychiatry first; "
        "decision in ~77 days. If rejected, submit to Asian Journal of Psychiatry (best "
        "cultural fit).",
        "Tier 2 — Springer/Wiley backup (weeks 12–24). If Tier 1 rejects, submit to Child "
        "Psychiatry & Human Development or School Mental Health (both Springer ecosystem).",
        "Tier 3 — open-access broad (weeks 24–36). If Tier 2 rejects, BMC Psychology "
        "(€1,390 — cheapest) or Frontiers in Psychology. Cambridge Prisms: Global Mental "
        "Health is particularly suitable due to its LMIC focus.",
        "Tier 4 — policy/services or Q3 (weeks 36+). If still rejected, Children and Youth "
        "Services Review (policy lane), Journal of School Health (school-based emphasis), or "
        "Vulnerable Children and Youth Studies (safe Q3 fallback).",
    ]
    for c in cascade:
        add_para(doc, "• " + c, size=12, indent=0.3, align="justify")

    add_h2(doc, "Warnings to avoid")
    warns = [
        "Avoid MDPI International Journal of Environmental Research and Public Health "
        "(IJERPH) — Clarivate de-indexed this journal from Web of Science in 2024 due to "
        "citation-inflation concerns.",
        "Use caution with hybrid (subscription + OA) journals — read APC policies carefully.",
        "Avoid journals with acceptance rates below 10% unless the manuscript has been "
        "professionally English-edited.",
    ]
    for w in warns:
        add_para(doc, "• " + w, size=12, indent=0.3, align="justify")

    add_h2(doc, "Submission preparation notes")
    notes = [
        "All Q2–Q3 journals require a short cover letter (≤ 1 page) highlighting novelty and "
        "scope fit.",
        "The manuscript should be edited by a native English speaker (Editage, Enago, "
        "American Journal Experts) before Tier 1 submission, since Frontiers / Springer / "
        "Elsevier journals may desk-reject on language grounds.",
        "Declare AI tool (LLM) usage in the cover letter and Methods section, in line with "
        "the 2024 policies of major publishers (Elsevier, Springer, Wiley, Frontiers).",
        "Explicitly cite the three companion papers (Công and Đào, 2026 validation; two VJES "
        "narrative reviews) in Introduction and Methods to avoid editor concerns about salami "
        "slicing.",
    ]
    for n in notes:
        add_para(doc, "• " + n, size=12, indent=0.3, align="justify")


def main():
    print("Adding journal list to VN...")
    doc = Document(VN)
    add_vn_section(doc)
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""
    cp.subject = ""; cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 5, 1, 9, 0, 0)
    cp.modified = datetime(2026, 5, 15, 9, 0, 0)
    doc.save(VN)
    print(f"  Saved: {VN}")

    print("Adding journal list to EN...")
    doc = Document(EN)
    add_en_section(doc)
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""
    cp.subject = ""; cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 5, 1, 9, 0, 0)
    cp.modified = datetime(2026, 5, 15, 9, 0, 0)
    doc.save(EN)
    print(f"  Saved: {EN}")
    print("\n[DONE]")


if __name__ == "__main__":
    main()
