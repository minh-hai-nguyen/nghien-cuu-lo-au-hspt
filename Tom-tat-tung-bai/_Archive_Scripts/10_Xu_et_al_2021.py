# -*- coding: utf-8 -*-
"""Bài 10: Xu et al. 2021 — CTH v5 (≤2,5 trang)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.5
for s in doc.sections: s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)

def rb(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

bl('Tóm tắt bài 10', bold=True)

rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình \u00ab Tỷ lệ và các yếu tố nguy cơ đối với triệu chứng lo âu trong thời kỳ bùng phát COVID-19: Khảo sát quy mô lớn trên 373.216 học sinh trung học tại Trung Quốc \u00bb (Prevalence and Risk Factors for Anxiety Symptoms during the Outbreak of COVID-19), do Qingqing Xu và cộng sự (2021), khảo sát trên 373.216 học sinh (244.193 THCS + 129.023 THPT, tuổi TB 15,24) tại 3 thành phố tỉnh Hà Nam, Trung Quốc, trong 8 ngày (4\u201312/02/2020 \u2014 đỉnh dịch COVID-19), xuất bản trên Journal of Affective Disorders, 288:17\u201322 (tạp chí Q1).')

rb('Phương pháp nghiên cứu')
bl('Công trình sử dụng GAD-7 (Thang đo Rối loạn Lo âu Tổng quát, ngưỡng ≥10 cho dương tính) kết hợp đánh giá kiến thức COVID-19 (3 mục), mức lo lắng (Likert 5 điểm) và hành vi phòng ngừa (4 mục). Nói cách khác, đây là khảo sát cắt ngang trực tuyến quy mô cực lớn \u2014 CỠ MẪU LỚN NHẤT TOÀN CẦU về lo âu thanh thiếu niên.', bold=True)

rb('Kết quả nghiên cứu định lượng')
tbl(
    ['Nhóm', 'Tỷ lệ lo âu (GAD-7 ≥10)', 'Ghi chú'],
    [
        ['Tổng thể', '9,89% (36.918/373.216)', '\u2014'],
        ['THCS', '10,85%', 'Cao hơn'],
        ['THPT', '8,08%', 'Thấp hơn'],
        ['Nam', '10,11%', 'Cao hơn nữ'],
        ['Nữ', '9,66%', '\u2014'],
        ['Nông thôn', '11,33%', 'Cao nhất'],
        ['Thành thị', '8,77%', 'Thấp nhất'],
    ],
    widths=[3.5, 4.5, 3.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan (hồi quy logistic đa biến, THCS):', bold=True)
tbl(
    ['Yếu tố', 'OR', 'KTC 95%'],
    [
        ['Nông thôn (vs thành thị)', '1,30', '1,26\u20131,34'],
        ['Hành vi phòng ngừa kém', '2,72', '2,01\u20133,68'],
        ['Nữ (vs nam)', '0,92 (bảo vệ)', '0,89\u20130,94'],
        ['Lo lắng COVID TB (vs cao)', '0,60', '0,56\u20130,64'],
        ['Sợ hãi TB (vs cao)', '0,21', '0,20\u20130,22'],
    ],
    widths=[4.5, 2.5, 4.0]
)
doc.add_paragraph()

bh3('Khác biệt giới tính : Trái ngược với đa số y văn, nam có tỷ lệ lo âu cao hơn nữ (10,11% vs 9,66%, P < 0,001). Trong 11 bài NC của Đề tài, chỉ có 2 nghiên cứu cho thấy nam > nữ: Saikia 2023 (Đông Bắc Ấn Độ) và Xu 2021 (bài này). Điều này có thể do bối cảnh COVID-19 và đặc thù nông thôn.')

rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bl('*Cỡ mẫu kỷ lục.* 373.216 học sinh là cỡ mẫu lớn nhất toàn cầu về lo âu thanh thiếu niên, đảm bảo sức mạnh thống kê rất cao. Tỷ lệ 9,89% sử dụng ngưỡng GAD-7 ≥10 (trung bình-nặng) là bảo thủ so với ngưỡng ≥5.')
bl('*Nông thôn > thành thị.* Khoảng cách nông thôn-thành thị (11,33% vs 8,77%, OR = 1,30) phản ánh bất bình đẳng trong tiếp cận dịch vụ SKTT, tương tự phát hiện của Nakie (2022) tại Ethiopia.')

rb('Kết luận')
bl('Tỷ lệ lo âu 9,89% trong đỉnh dịch COVID-19 ở 373.216 học sinh Trung Quốc. Nông thôn, THCS và hành vi phòng ngừa kém là yếu tố nguy cơ. Sợ hãi quá mức về COVID-19 liên quan mạnh đến lo âu (OR sợ hãi cao gấp 5 lần). Can thiệp cần giảm sợ hãi không hợp lý và tăng cường dịch vụ SKTT tại vùng nông thôn.', bold=True)

rh2('Phản biện')
bl('J Affect Disord Q1, cỡ mẫu kỷ lục. Tuy nhiên, khảo sát trực tuyến trong 8 ngày có thiên lệch tự chọn. Chỉ dùng GAD-7 (sàng lọc, không chẩn đoán). Thu thập đỉnh dịch (02/2020) \u2014 kết quả phản ánh bối cảnh khủng hoảng cấp tính, không thể khái quát cho thời bình. Kiến thức COVID-19 đánh giá bằng 3 câu chưa xác thực.')

rh2('Hướng nghiên cứu tiếp theo')
bl('So sánh dọc lo âu đỉnh dịch vs hậu COVID. Nghiên cứu can thiệp giảm sợ hãi không hợp lý ở thanh thiếu niên. Mở rộng đến nhiều tỉnh ngoài Hà Nam.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50 Cao. Q1, cỡ mẫu kỷ lục 373.216, GAD-7 xác thực, nhưng bối cảnh COVID-19 hạn chế tính khái quát.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('10_Xu_et_al_2021.docx')
print('10_Xu_et_al_2021.docx OK')
