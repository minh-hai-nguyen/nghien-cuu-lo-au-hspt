# -*- coding: utf-8 -*-
"""Dump full Alharbi tom tat to check numbers."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
F = os.path.join(ROOT, 'Tom-tat-tung-bai/QT005_Alharbi_et_al_2019_SaudiArabia.docx')
d = Document(F)
print("=== Alharbi tom tat - all paragraphs ===")
for i, p in enumerate(d.paragraphs):
    if p.text.strip():
        print(f"  [{i}] {p.text[:400]}")
print(f"\n=== Tables ===")
for ti, tb in enumerate(d.tables):
    print(f"\n-- Table {ti+1} --")
    for ri, row in enumerate(tb.rows):
        print(f"  Row {ri}: {[c.text[:80] for c in row.cells]}")
