"""DOC: Tra loi cau hoi cua thay:
'Khong phai o LAXH, nu co muc do khac biet cao hon nam so voi LATQ?'

Hai goc nhin: DTB raw + Cohen d
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Chenh_lech_gioi_LAXH_vs_LATQ_2_goc_nhin.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CHÊNH LỆCH GIỚI Ở LO ÂU XÃ HỘI vs LO ÂU LAN TỎA\n— Hai góc nhìn cho cùng câu hỏi —\nTrả lời thầy ngày 09/05/2026')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue(
    'Không phải là ở LAXH, nữ có mức độ khác biệt cao hơn nam '
    'so với LATQ?', italic=True
)

# Trả lời ngắn
H('Trả lời NGẮN — Thầy ĐÚNG theo MỘT GÓC NHÌN', level=2, color=GREEN)
para(
    'Theo CHÊNH LỆCH ĐTB RAW (điểm tuyệt đối trên thang 0–100), '
    'thầy ĐÚNG: ở LAXH chênh lệch nữ-nam = 9,54 điểm, lớn hơn '
    'ở LATQ (8,04 điểm) khoảng 18,7%. Tuy nhiên, theo Cohen d '
    '(effect size chuẩn hóa theo SD), chênh lệch giới ở hai '
    'dạng gần như BẰNG NHAU (d = 0,370 vs 0,365 — chênh 1,5%).', bold=True
)

# I. Hai góc nhìn
H('I. Hai góc nhìn cho cùng câu hỏi', level=2, color=NAVY)

H('Góc nhìn 1 — ĐTB RAW (chênh lệch điểm tuyệt đối): THẦY ĐÚNG', level=3, color=GREEN)
add_table(
    ['Dạng', 'Nam ĐTB', 'Nữ ĐTB', 'Chênh nữ−nam'],
    [
        ['LATQ (lan tỏa)', '51,43', '59,47', '8,04 điểm'],
        ['LAXH (xã hội)', '43,20', '52,74', '9,54 điểm'],
    ]
)
para('')
para(
    'Ở LAXH, nữ vượt nam 9,54 điểm — LỚN HƠN ở LATQ (8,04 điểm) '
    'khoảng 1,5 điểm tuyệt đối, hay +18,7% tỷ lệ. Trên thang '
    'điểm thực tế, đây là KHÁC BIỆT đáng kể.', bold=True
)

H('Góc nhìn 2 — Cohen d (chuẩn hóa theo SD): GẦN BẰNG NHAU', level=3, color=NAVY)
add_table(
    ['Dạng', 'Chênh raw', 'SD pooled', 'Cohen d = chênh / SD'],
    [
        ['LATQ', '8,04', '22,04', '0,365'],
        ['LAXH', '9,54', '25,77', '0,370'],
    ]
)
para('')
para(
    'Sau khi chuẩn hóa theo SD, chỉ chênh +1,5% — Cohen d hai '
    'dạng GẦN BẰNG NHAU.', bold=True
)

# II. Tại sao 2 góc nhìn khác nhau
H('II. Vì sao hai góc nhìn cho kết quả khác nhau?', level=2, color=NAVY)
para(
    'Lý do CHÍNH: SD ở LAXH (25,77) LỚN HƠN SD ở LATQ (22,04). '
    'Có nghĩa các học sinh ở LAXH PHÂN TÁN HƠN quanh ĐTB so với '
    'LATQ.'
)
para('')
add_table(
    ['Đặc điểm', 'LATQ', 'LAXH'],
    [
        ['Nam: ĐTB ± SD', '51,43 ± 22,01', '43,20 ± 25,09'],
        ['Khoảng nam (~68% mẫu)', '29 → 73', '18 → 68'],
        ['Nữ: ĐTB ± SD', '59,47 ± 22,07', '52,74 ± 26,31'],
        ['Khoảng nữ (~68% mẫu)', '37 → 81', '26 → 79'],
        ['Chồng lấn nam-nữ', 'Vừa phải', 'Nhiều hơn (SD lớn)'],
    ]
)
para('')
para(
    'Diễn giải đơn giản: ở LAXH, dù chênh ĐTB nữ-nam LỚN HƠN, '
    'các điểm cá nhân CHỒNG LẤN nhiều hơn vì biến thiên trong '
    'nhóm cao. Có học sinh nam có điểm rất cao (sợ xã hội), có '
    'học sinh nữ có điểm rất thấp (bình thản). Cohen d tính '
    'đến điều này nên kết quả "co lại" tương đối.'
)
para(
    'Ý nghĩa thực tế: lo âu xã hội PHÂN HÓA mạnh hơn lo âu lan '
    'tỏa — một số HS rất sợ tình huống xã hội, số khác bình '
    'thản; trong khi lo âu lan tỏa phân bố đều hơn.'
)

# III. Diễn giải đúng cho báo cáo
H('III. Diễn giải đúng cho báo cáo CTH', level=2, color=NAVY)
para('Cả hai cách đều ĐÚNG — nên báo cáo SONG HÀNH:', bold=True)
para('')
blue(
    '"Trên thang điểm thực tế (0–100), chênh lệch giới ở rối loạn '
    'lo âu xã hội là 9,54 điểm — LỚN HƠN ở rối loạn lo âu lan '
    'tỏa (8,04 điểm) khoảng 18,7%. Tuy nhiên, sau khi chuẩn hóa '
    'theo độ lệch chuẩn (Cohen d), chênh lệch giới ở hai dạng '
    'gần như TƯƠNG ĐƯƠNG (d = 0,370 và 0,365 — chênh 1,5%). Điều '
    'này phản ánh thực tế: lo âu xã hội có biến thiên cá nhân '
    'lớn hơn (SD pooled = 25,77 so với 22,04 ở lan tỏa) — một '
    'số học sinh rất sợ tình huống xã hội, một số bình thản; '
    'trong khi lo âu lan tỏa phân bố đều hơn."', italic=True
)

# IV. So sánh ba góc nhìn
H('IV. Bảng tổng so sánh ba chỉ số', level=2, color=NAVY)
add_table(
    ['Chỉ số', 'LATQ', 'LAXH', 'XH lớn hơn TQ', 'Đánh giá'],
    [
        ['Chênh ĐTB raw (điểm)', '8,04', '9,54', '+1,5 điểm (+18,7%)', 'LỚN HƠN ĐÁNG KỂ'],
        ['F-test (ANOVA)', '44,484', '45,984', '+1,5 (+3,4%)', 'Nhỏ — không đáng kể'],
        ['Cohen d (chuẩn hóa)', '0,365', '0,370', '+0,005 (+1,5%)', 'Cực nhỏ — gần bằng'],
    ]
)

# V. Khuyến nghị cách phát biểu
H('V. Khuyến nghị cách phát biểu', level=2, color=NAVY)

para('CÁCH PHÁT BIỂU TỐT NHẤT (kết hợp cả 2 góc nhìn):', bold=True, color=GREEN)
blue(
    '"Cả lo âu lan tỏa và lo âu xã hội đều có chênh lệch giới '
    'có ý nghĩa thống kê (p < 0,001 cho cả hai), với nữ vượt '
    'nam. Trên thang điểm thực tế, chênh lệch lớn hơn ở lo âu '
    'xã hội (9,54 điểm vs 8,04 điểm — chênh 18,7%), tuy nhiên '
    'sau chuẩn hóa effect size theo Cohen d, hai dạng có cường '
    'độ tương đương (d ≈ 0,37 cho cả hai, small-to-medium). '
    'Nguyên nhân: lo âu xã hội có biến thiên cá nhân lớn hơn '
    'lo âu lan tỏa."'
)

para('')
para('TRÁNH các phát biểu ĐƠN PHƯƠNG:', bold=True)
para('')
para('❌ "Chênh lệch giới ở LAXH lớn hơn LATQ" — đúng raw nhưng không đúng effect size, có thể gây hiểu nhầm.')
para('❌ "Chênh lệch giới ở LAXH và LATQ không khác nhau" — đúng effect size nhưng bỏ qua thông tin raw có giá trị.')
para('')
para('✓ "Chênh lệch giới ở LAXH lớn hơn LATQ về điểm tuyệt đối (+19%) nhưng tương đương về effect size chuẩn hóa." — Đầy đủ và chính xác.', bold=True)

# VI. Tóm gọn xanh
H('VI. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue('Thầy ĐÚNG ở góc nhìn ĐTB RAW:', bold=True)
blue('• Ở LAXH: chênh nữ-nam = 9,54 điểm')
blue('• Ở LATQ: chênh nữ-nam = 8,04 điểm')
blue('• LAXH lớn hơn LATQ: +1,5 điểm tuyệt đối, +18,7% tỷ lệ')
blue('')
blue('Tuy nhiên ở góc nhìn EFFECT SIZE chuẩn (Cohen d):', bold=True)
blue('• d LATQ = 0,365')
blue('• d LAXH = 0,370')
blue('• Chênh chỉ 1,5% — gần như bằng nhau')
blue('')
blue('Lý do khác biệt: SD của LAXH (25,77) lớn hơn SD của LATQ '
     '(22,04). Khi chuẩn hóa theo SD, chênh raw bị "co lại" '
     'tương đối.', bold=True)
blue('')
blue('Khuyến nghị: báo cáo CẢ HAI để đầy đủ và chính xác. Hội '
     'đồng VN có thể nhìn theo cả 2 góc — không có cách nào '
     '"sai", chỉ TRẢ LỜI CÂU HỎI KHÁC nhau.', bold=True)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
