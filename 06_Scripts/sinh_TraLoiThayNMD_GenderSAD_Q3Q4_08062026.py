# -*- coding: utf-8 -*-
"""Doc tra loi thay NMD 08/06/2026:
- Phan tich gender difference SAD (HS THCS VN vs Jefferies 16-29)
- De xuat 3 option Q3 + 3 option Q4."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'TraLoiThayNMD_GenderSAD_Q3Q4_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def H1(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H3(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(t, italic=False, indent=True):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = italic

def B(t, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

def REC(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('► ĐỀ XUẤT: ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)


# ============================================================
H1('TRẢ LỜI THẦY NGUYỄN MINH ĐỨC')
P('Phân tích gender difference SAD + Đề xuất tiêu đề Q3 và Q4',
  italic=True, indent=False)
P('Soạn ngày 08/06/2026', italic=True, indent=False)


# ============================================================
H2('PHẦN A — XÁC NHẬN BA NỘI DUNG THẦY GIAO')

P('Em xin cám ơn thầy Nguyễn Minh Đức đã đánh giá tích cực bản thảo Q2 '
  'và ghi nhận hàm lượng khoa học gắn với văn hóa Á Đông – Việt Nam của '
  'bài. Em xin xác nhận ba quyết định nhóm và ba việc cần làm thêm sau '
  'tin nhắn của thầy:')

B('Quyết định 1: Nhóm đồng ý chọn Frontiers in Psychiatry làm '
  'target Q2 chính (chấp nhận IF thấp hơn để đổi lấy thời gian peer-'
  'review nhanh + an toàn). Em xin lưu ý nhóm một cảnh báo về dòng tạp '
  'chí Frontiers em đã ghi rõ trong PhuongAnXuLy mục 2.3.', 0)

B('Quyết định 2: NCS Công Thị Hằng sẽ làm việc với cơ sở đào tạo '
  '(HNUE) để giải quyết thủ tục chấp thuận của Hội đồng đạo đức — giải '
  'quyết Q3-6 BLOCKING.', 0)

B('Việc bổ sung 1: Em sẽ bổ sung phần bình luận/luận chứng vào Q2 '
  'v1 để lý giải gender difference SAD ở HS THCS Việt Nam vs Jefferies '
  '& Ungar 2020 — trình bày trong PHẦN B dưới đây.', 0)

B('Việc bổ sung 2: Em đề xuất tiêu đề Q3 + Q4 tiềm năng — trình bày '
  'trong PHẦN C.', 0)


# ============================================================
H2('PHẦN A bổ sung — Ghi nhận đánh giá của thầy về hàm lượng khoa học '
   'Q2 v1 so với Jefferies & Ungar 2020')

P('Em cám ơn thầy đã ghi nhận hàm lượng khoa học gắn với văn hóa Á '
  'Đông và Việt Nam của Q2 v1. Em xin tóm lược năm khía cạnh để hai '
  'nghiên cứu bổ sung lẫn nhau hơn là cạnh tranh:')

B('Phương pháp phân tích — Q2 v1 sử dụng SEM tích hợp với mô hình '
  'bậc cao YTNC + YTBV và kiểm định bất biến đa nhóm theo giới '
  '(configural → metric → scalar), trong khi Jefferies dùng ANOVA + '
  't-test + chi-square (mô tả). Hai phương pháp có mức độ chặt chẽ '
  'thống kê khác nhau và phù hợp với câu hỏi nghiên cứu khác nhau.', 0)

B('Khung lý thuyết văn hóa — Q2 v1 áp dụng ba khung lý thuyết Á '
  'Đông cụ thể: tam giáo Nho-Phật-Đạo (Small & Blanc 2021), văn hóa '
  'học thuật Nho giáo (Stankov 2010), co-rumination ở vị thành niên '
  '(Rose 2002). Jefferies dừng ở phân tích descriptive country '
  'effects — không khai thác sâu vào cơ chế văn hóa.', 0)

B('Bối cảnh phát triển — Q2 v1 tập trung vào cohort phát triển đồng '
  'nhất (HS THCS 11–14 tuổi, mọi đối tượng đang trong cùng giai '
  'đoạn dậy thì và xã hội hóa). Jefferies pool ba giai đoạn phát '
  'triển khác biệt (16–17 tuổi học sinh phổ thông, 18–24 tuổi sinh '
  'viên đại học, 25–29 tuổi đã đi làm) — phân tích pooled có thể '
  'che giấu các hiệu ứng đặc thù từng giai đoạn.', 0)

B('Thiết kế lấy mẫu — Q2 v1 sử dụng khảo sát học đường với purposive '
  'sampling 2 trường tại Hà Nội (đại diện đô thị + ngoại ô), gắn với '
  'bối cảnh giáo dục cụ thể. Jefferies là secondary analysis trên '
  'dataset thu thập bởi công ty market research Edelman/Unilever cho '
  'mục đích nghiên cứu thị trường mỹ phẩm — không gắn với bối cảnh '
  'lâm sàng hay giáo dục cụ thể nào.', 0)

B('Công cụ đo — Q2 v1 sử dụng RCADS-SocAD (DSM-5 aligned, 4 item '
  'tối ưu cho 11–14 tuổi). Jefferies sử dụng SIAS 20 item (Social '
  'Interaction Anxiety Scale, không DSM-5 aligned cụ thể). Hai công '
  'cụ đo cấu trúc gần kề nhưng không hoàn toàn trùng — RCADS-SocAD '
  'tập trung vào tiêu chí chẩn đoán SAD, SIAS tập trung vào tương '
  'tác xã hội tổng quát.', 0)

P('Điểm mạnh của Jefferies là cỡ mẫu lớn (N = 6.825) và phạm vi đa '
  'quốc gia (7 nước) — Q2 v1 không tham vọng cạnh tranh ở hai khía '
  'cạnh này. Q2 v1 đóng góp bổ sung ở chiều sâu cơ chế tâm lý-văn hóa '
  'và phương pháp luận SEM tích hợp cho một cohort phát triển đồng '
  'nhất tại một quốc gia Á Đông cụ thể.')


# ============================================================
H2('PHẦN B — LUẬN CHỨNG VỀ GENDER DIFFERENCE SAD')

H3('B1. So sánh hai nghiên cứu')

P('Em đã đọc lại Jefferies & Ungar (2020) đầy đủ. So sánh hai nghiên cứu '
  'như sau:')

t = d.add_table(rows=1, cols=3); t.style = 'Light Grid Accent 1'
t.autofit = False
hdr = t.rows[0].cells
for i, h in enumerate(['Tiêu chí',
                       'Jefferies & Ungar 2020 [1]',
                       'Q2 v1 của NCS']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)

rows = [
    ['Cỡ mẫu', 'N = 6.825', 'N = 1.352'],
    ['Độ tuổi', '16–29 (M = 22,84; SD = 3,97)', '11–14 (HS THCS)'],
    ['Giới', '3.342 nam / 3.428 nữ', '614 nam / 738 nữ'],
    ['Phạm vi', '7 quốc gia (Brazil, China, Indonesia, Russia, '
                'Thailand, US, Vietnam)', 'Việt Nam (Hà Nội)'],
    ['Nguồn mẫu', 'Online market research panel (Edelman/Unilever); '
                  '23.346 invited → 6.825 final', 'Khảo sát trường học '
                  '(Nhật Tân + Tây Mỗ)'],
    ['Công cụ đo', 'SIAS (Social Interaction Anxiety Scale, 20 item; '
                   'em đã verify trong PDF)', 'RCADS-SocAD (4 item, '
                   'DSM-5 aligned)'],
    ['Phát hiện gender chính', 't(6768) = −1,37; n.s. — '
                              'KHÔNG có khác biệt giới',
     'Nữ > nam có ý nghĩa thống kê'],
    ['Phát hiện country × sex', 'F(6,6728) = 2,25; p = 0,036; '
                                'η² = 0,002 (có ý nghĩa, '
                                'effect size rất nhỏ)',
     'Không tham chiếu — single country'],
    ['Phát hiện age main effect', 'F(2,6822) = 39,74; p < 0,001; '
                                  '18–24 cao nhất (M = 25,33), '
                                  '16–17 thấp nhất (M = 21,92), '
                                  '25–29 ở giữa (M = 22,44)',
     'Không kiểm — sample 11–14 đồng nhất'],
    ['Mẫu Việt Nam sub-sample', 'N = 984 (487 nam / 493 nữ); '
                                'M = 22,68; SD = 11,77 — xếp thứ 4 '
                                'trong 7 quốc gia (Indonesia 18,94 < '
                                'Nga 20,78 < Trung Quốc 22,30 < '
                                'Việt Nam 22,68), dưới mức trung bình '
                                'toàn mẫu 23,82',
     '— (đã là toàn bộ mẫu)'],
]
for row_data in rows:
    row = t.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
for row in t.rows:
    widths = [3.5, 7.0, 6.0]
    for idx, w in enumerate(widths):
        if idx < len(row.cells):
            row.cells[idx].width = Cm(w)

P('Câu hỏi của thầy thực sự rất sắc bén — vì nếu chỉ nhìn vào kết quả '
  'tổng thể của Jefferies (gender effect n.s.) mà chưa cân nhắc các '
  'điều kiện biên (boundary conditions), người đọc dễ kết luận rằng '
  'gender difference SAD không tồn tại, do đó kết quả của NCS có vẻ '
  'mâu thuẫn với y văn quốc tế. Em xin đề xuất 6 luận chứng — chia '
  'thành hai nhóm rõ rệt — để giải thích sự không nhất quán này:')

B('Bốn luận chứng đầu (1–4) giải thích vì sao gender difference '
  'XUẤT HIỆN trong mẫu HS THCS Việt Nam — đây là các cơ chế tâm lý-'
  'phát triển-văn hóa làm cho nữ sinh 11–14 tuổi đặc biệt dễ tổn '
  'thương SAD hơn nam.', 0)

B('Hai luận chứng cuối (5–6) giải thích vì sao gender difference '
  'KHÔNG XUẤT HIỆN trong mẫu Jefferies — đây là các yếu tố phương '
  'pháp luận làm cho khác biệt giới khó được phát hiện dù có thể tồn '
  'tại trong dân số.', 0)

H3('B2. Sáu luận chứng đề xuất')

H3('Luận chứng 1 — Cửa sổ phát triển 11–14 tuổi là đỉnh điểm nguy cơ SAD ở '
   'nữ')

P('Tổng quan Hankin và cộng sự (2007) [2] trên cohort thanh thiếu niên '
  'Mỹ cho thấy khác biệt giới về triệu chứng nội hóa (depression, '
  'anxiety) bắt đầu xuất hiện rõ rệt vào khoảng 13 tuổi và đạt đỉnh ở '
  '15–17 tuổi. McLean và Anderson (2009) [3] trong tổng quan về khác '
  'biệt giới trong fear và anxiety chỉ ra rằng SAD cụ thể có xu hướng '
  'thể hiện sớm hơn ở nữ do trải nghiệm hội tụ ba yếu tố vào đầu vị '
  'thành niên: (a) thay đổi nội tiết do dậy thì (nữ bước vào dậy thì '
  'sớm hơn nam khoảng 2 năm), (b) tăng nhạy cảm với đánh giá của bạn '
  'bè đồng trang lứa do trưởng thành thần kinh vỏ não trán-rìa, '
  '(c) nội hóa vai trò giới về ngoại hình + đánh giá xã hội theo lý '
  'thuyết gender intensification (Hill & Lynch 1983) [4]. Theo Kessler '
  'và cộng sự (2005) [5] từ National Comorbidity Survey Replication, '
  'một nửa số ca rối loạn lo âu khởi phát trước 14 tuổi — gợi ý rằng '
  'cửa sổ 11–14 chính là giai đoạn dễ tổn thương nhất. Đến 16–29 tuổi '
  'trong mẫu Jefferies, nam đã trải qua cùng tiến trình dậy thì (chậm '
  'hơn 2 năm) và đã đạt mức tổn thương tương đương, làm thu hẹp khoảng '
  'cách giới đo lường được.')

H3('Luận chứng 2 — Co-rumination đạt đỉnh ở giữa vị thành niên, đặc biệt '
   'ở nữ (Rose 2002)')

P('Bên cạnh cơ chế phát triển sinh học và xã hội-văn hóa nói chung, '
  'có một dạng tương tác xã hội đặc trưng của nữ sinh trung học cơ sở '
  'càng làm khuếch đại nguy cơ SAD:', italic=True)

P('Em đã sử dụng co-rumination hypothesis của Rose (2002) [6] trong '
  'Discussion Q2 v1 để giải thích kết quả null của peer support. Khái '
  'niệm này cũng giải thích trực tiếp gender difference SAD: trong mẫu '
  'lớp 3, 5, 7, 9 (Rose 2002, N = 608), nữ sinh trung học cơ sở có tần '
  'suất co-rumination cao nhất so với cả nam cùng tuổi và nữ ở cấp cao '
  'hơn. Tăng co-rumination → khuếch đại lo lắng → tăng triệu chứng SAD. '
  'Đến tuổi trưởng thành, mẫu hình co-rumination ổn định lại và khoảng '
  'cách giới hẹp dần.')

H3('Luận chứng 3 — Kịch bản giới văn hóa Nho giáo Việt Nam đặc biệt '
   'nghiêm ngặt ở tuổi học sinh THCS')

P('Hai cơ chế trên (phát triển sinh học + co-rumination) mang tính phổ '
  'quát qua các nền văn hóa. Bối cảnh Việt Nam còn áp đặt thêm một lớp '
  'áp lực giới đặc thù làm khuếch đại lo âu xã hội ở nữ:', italic=True)

P('Stankov (2010) [7] đã chỉ ra rằng văn hóa học thuật Nho giáo nuôi '
  'dưỡng test anxiety + self-doubt ở học sinh Đông Á. Trong bối cảnh '
  'Việt Nam cụ thể, nữ sinh THCS chịu kép áp lực: vừa phải đạt thành '
  'tích học tập cao (theo Confucian achievement script), vừa phải đáp '
  'ứng kỳ vọng giới về tế nhị, kín đáo, thu mình trong giao tiếp xã hội '
  '(theo Confucian feminine script). Nam sinh cùng tuổi chịu áp lực '
  'thành tích nhưng được phép biểu hiện chủ động xã hội. Small & Blanc '
  '(2021) [8] mô tả tam giáo Nho-Phật-Đạo định hình các chuẩn mực ứng '
  'xử kéo dài qua nhiều thế hệ; những kỳ vọng giới này đặc biệt nghiêm '
  'ngặt trong môi trường học đường dưới sự giám sát của giáo viên + '
  'gia đình. Đến 16–29 tuổi (mẫu Jefferies), một bộ phận nữ thanh '
  'niên đã rời ghế nhà trường, vào môi trường lao động độc lập hơn — '
  'áp lực giới giảm xuống.')

H3('Luận chứng 4 — Bối cảnh đo lường khác nhau: học đường vs population-'
   'based panel')

P('Áp lực giới mà luận chứng 3 mô tả biểu hiện rõ nhất ở môi trường '
  'có sự giám sát xã hội cao — đây cũng là môi trường mà Q2 v1 chọn '
  'để đo lường:', italic=True)

P('Q2 v1 đo lo âu xã hội trong môi trường mà các tình huống đánh giá '
  'xã hội xảy ra hàng ngày: phát biểu trước lớp, kiểm tra miệng, '
  'tương tác giáo viên, xếp hạng học lực công khai. Đây là môi trường '
  '"nuôi dưỡng" SAD đặc biệt cho nữ sinh có xu hướng nội hóa đánh giá '
  'tiêu cực. Trong khi đó, Jefferies sử dụng online market research '
  'panel không gắn với bối cảnh xã hội cụ thể nào — đo lường mức lo '
  'âu tổng quát trong cuộc sống hàng ngày. Hai bối cảnh đo lường khác '
  'nhau có thể cho kết quả khác nhau ngay cả khi đo cùng một construct.')

H3('Luận chứng 5 — Khác biệt phương pháp luận: cỡ mẫu + công cụ + sai '
   'lệch mẫu trong nghiên cứu Jefferies')

P('Bốn luận chứng trên (1–4) tập trung vào cơ chế tâm lý-văn hóa giải '
  'thích vì sao gender difference XUẤT HIỆN ở mẫu Q2 v1. Hai luận '
  'chứng tiếp theo (5–6) chuyển sang phía Jefferies để giải thích vì '
  'sao gender difference KHÔNG ĐƯỢC PHÁT HIỆN ở mẫu của họ — đây là '
  'các yếu tố phương pháp luận quan trọng:', italic=True)

P('Năm khía cạnh phương pháp luận góp phần giải thích:', italic=False)

B('Công cụ khác nhau: Jefferies dùng SIAS 17–20 item đo Social '
  'INTERACTION Anxiety (tương tác xã hội tổng quát); Q2 v1 dùng '
  'RCADS-SocAD 4 item đo Social Anxiety DISORDER theo tiêu chí DSM-5. '
  'Hai công cụ đo cấu trúc khác nhau và có thể nhạy cảm với gender '
  'difference khác nhau.', 1)

B('Sai lệch lựa chọn của Jefferies: 23.346 invited → 6.825 final '
  '(76% panel members consented; 65% rejected for eligibility hoặc '
  'occupation; 176 loại bỏ vì straight-lining). Người có SAD nặng '
  'thường tránh khảo sát online interactive → underrepresented trong '
  'mẫu. Đặc biệt nữ SAD nặng có thể bị mất hơn nam → che giấu sự khác '
  'biệt giới thực sự. Jefferies thừa nhận hạn chế này.', 1)

B('Cỡ mẫu thực tế cho phân tích Việt Nam: Jefferies có N = 984 cho '
  'Việt Nam (487 nam / 493 nữ) — đây là cỡ mẫu khá nhỏ khi xét phân '
  'tích con cho từng quốc gia. Khả năng phát hiện gender difference '
  'cấp quốc gia thấp hơn so với Q2 v1 (N = 1.352 trong cùng quốc gia).', 1)

B('Country × sex interaction CÓ ý nghĩa thống kê trong Jefferies '
  '(F = 2,25; p = 0,036) — nghĩa là gender effect KHÁC NHAU theo '
  'quốc gia. Đây là bằng chứng quan trọng: Jefferies KHÔNG bác bỏ '
  'gender difference theo quốc gia; chỉ kết luận rằng main effect '
  'pooled qua 7 quốc gia là n.s.', 1)

B('Mẫu Jefferies cho Việt Nam có M = 22,68 — xếp thứ tư từ thấp lên '
  'trong 7 quốc gia (lần lượt: Indonesia 18,94 < Nga 20,78 < Trung '
  'Quốc 22,30 < Việt Nam 22,68), thấp hơn mức trung bình toàn mẫu '
  '23,82. Khả năng mẫu Việt Nam của Jefferies là một low-anxiety '
  'subgroup do tự chọn lọc participants — điều này làm giảm khả năng '
  'phát hiện gender difference vốn có thể tồn tại trong dân số thanh '
  'thiếu niên Việt Nam nói chung.', 1)

H3('Luận chứng 6 — Hiệu ứng tuổi mạnh áp đảo trong Jefferies + quá độ '
   'phát triển giữa hai mẫu')

P('Luận chứng cuối cùng — và có lẽ thuyết phục nhất — nằm chính trong '
  'kết quả của Jefferies:', italic=True)

P('Jefferies tự phát hiện chính trong nghiên cứu: F(2,6822) = 39,74; '
  'p < 0,001 — gender main effect không có nhưng AGE main effect rất '
  'mạnh. Phân tích chi tiết: 18–24 tuổi có lo âu cao nhất (M = 25,33), '
  '16–17 thấp nhất (M = 21,92), 25–29 ở giữa (M = 22,44). Mẫu nhỏ '
  'tuổi của Jefferies (16–17) cách mẫu Q2 v1 (11–14) khoảng 4–6 năm — '
  'một khoảng tuổi rất quan trọng trong quá trình ổn định gender role '
  'và mức lo âu xã hội. Hai nghiên cứu không thật sự đo cùng giai đoạn '
  'phát triển.')


# ============================================================
H2('PHẦN C — ĐỀ XUẤT TIÊU ĐỀ Q3 VÀ Q4')

H3('C1. Hai phương án cho lộ trình xuất bản')

P('Em đề xuất hai phương án theo gợi ý của thầy, tùy vào việc dữ liệu '
  'có đủ cho bài thứ ba (Q4) hay không:')

B('Phương án 2 bài (an toàn): Q2 + Q3 đồng hành submit song song; '
  'cả hai sử dụng chung dữ liệu định lượng 1.352 HS THCS Hà Nội.', 0)

B('Phương án 3 bài (mở rộng): Q2 + Q3 song hành như Phương án 2; Q4 '
  'là bài độc lập sử dụng phương pháp phân tích mới (latent class / '
  'network / framework can thiệp). Q4 chỉ khả thi nếu Q2 + Q3 đã thông '
  'qua review và dữ liệu cho phép tách phương pháp mới mà không trùng.', 0)

H3('C2. Ba đề xuất tiêu đề cho Q3 (companion với Q2)')

H3('Q3 — Đề xuất A: SEM phụ chuyên sâu về SAD')

REC('"Pathways to social anxiety disorder among Vietnamese lower secondary '
    'students: A subtype-specific structural equation analysis of school '
    'bullying victimization with developmental and Confucian cultural '
    'framing"')

P('Trọng tâm: đào sâu duy nhất phân loại SAD (vì β bắt nạt → SAD = 0,376 '
  'cao nhất trong Q2 v1) — phân tích mediation/moderation theo cấu '
  'trúc subtype-specific. Lợi thế: tách hoàn toàn về mặt trọng tâm với '
  'Q2 (Q2 = tích hợp 3 phân loại; Q3 = chuyên sâu SAD), không trùng '
  'lặp. Phù hợp scope Frontiers in Psychiatry section Anxiety and '
  'Stress Disorders. Liên kết tự nhiên với phân tích gender difference '
  'em vừa luận chứng trong PHẦN B.')

H3('Q3 — Đề xuất B: Multi-group SEM theo giới')

REC('"Gender-specific pathways to anxiety disorders in early adolescence: '
    'A multi-group structural equation model in Vietnamese lower secondary '
    'students with comparative discussion to cross-national young adult '
    'cohorts"')

P('Trọng tâm: phân tích bất biến đo lường và cấu trúc theo giới '
  '(configural → metric → scalar invariance, ΔCFI < 0,01) cho mô hình '
  'tích hợp YTNC + YTBV → 3 phân loại RLLA. Lợi thế: trực tiếp trả lời '
  'câu hỏi của thầy về cơ chế khác biệt giới, có novel theoretical '
  'contribution. So sánh với Jefferies 2020 + McLean 2011 làm framework. '
  'Phù hợp Frontiers Psychiatry section Adolescent Psychiatry hoặc '
  'BMC Public Health.')

H3('Q3 — Đề xuất C: Coping paradox + Confucian endurance script')

REC('"Adaptive intent, maladaptive enactment: The coping paradox among '
    'Vietnamese lower secondary students and the Confucian script of '
    'endurance"')

P('Trọng tâm: phân tích Brief COPE 3-yếu tố cùng RCADS — phát hiện '
  'counter-intuitive rằng cả problem-solving + support-seeking đều '
  'tương quan dương với lo âu (khác với expectation từ Lazarus & '
  'Folkman 1984 + Compas 2017). Giả thuyết frequency vs quality. '
  'Cultural framing tam giáo + tinh thần "nhẫn" như em đã verify với '
  'Small & Blanc 2021. Lợi thế: phát hiện ngược dòng, có thể attractive '
  'cho reviewer. Khó hơn: cần luận chứng theoretical sâu để defend '
  'paradox.')

H3('C3. Em đề xuất Q3 = Phương án B (gender-specific SEM)')

REC('Em đề xuất Q3 = Phương án B (Multi-group SEM theo giới) vì ba lý do.')

B('(1) Trực tiếp trả lời câu hỏi quan trọng của thầy về gender '
  'difference + dùng được hết các luận chứng em vừa soạn trong PHẦN B '
  '— không bỏ phí công sức.', 0)

B('(2) Có novel theoretical contribution (multi-group SEM trên cohort '
  '11–14 ở Việt Nam là một sự bổ sung hiếm trong y văn — em rà soát '
  'PubMed thấy gần như chưa có).', 0)

B('(3) Tách trọng tâm tốt với Q2 (Q2 = tích hợp, Q3 = phân nhóm giới). '
  'Có thể submit song song không sợ tự đạo văn.', 0)

H3('C4. Ba đề xuất tiêu đề cho Q4 (kế hoạch năm sau, nếu dữ liệu đủ)')

H3('Q4 — Đề xuất A: Latent profile analysis')

REC('"Anxiety phenotype profiles in Vietnamese lower secondary students: '
    'A latent profile analysis with risk-protective indicator integration"')

P('Phương pháp mới: LPA / LCA — phân loại học sinh thành các profile '
  'không quan sát được dựa trên kết hợp GAD + SAD + SocAD + YTNC + YTBV. '
  'Lợi thế: phương pháp khác Q2/Q3 (không phải SEM), không trùng lặp '
  'thống kê. Có thể publish ở tạp chí methodology-friendly như BMC '
  'Psychiatry hoặc Journal of Psychopathology and Behavioral Assessment.')

H3('Q4 — Đề xuất B: Network analysis')

REC('"Network structure of anxiety symptoms and risk factors among '
    'Vietnamese lower secondary students: A cross-sectional psychometric '
    'network analysis"')

P('Phương pháp mới: Gaussian Graphical Model + centrality analysis — '
  'phân tích cấu trúc mạng các triệu chứng lo âu + YTNC + YTBV thay vì '
  'mô hình hồi quy. Lợi thế: phương pháp hiện đại + mới được áp dụng '
  'cho Việt Nam. Phù hợp Frontiers in Psychology section Quantitative '
  'Psychology and Measurement.')

H3('Q4 — Đề xuất C: Theoretical framework can thiệp')

REC('"Designing subtype-specific prevention interventions for adolescent '
    'anxiety in Vietnam: A theoretical framework derived from differential '
    'pathway analysis"')

P('Bài lý thuyết / khuyến nghị chính sách: dựa trên kết quả Q2 + Q3 để '
  'đề xuất khung thiết kế can thiệp theo subtype-specific pathway. Lợi '
  'thế: ít cần dữ liệu mới — tổng hợp các kết quả + khuyến nghị. Phù '
  'hợp Health Policy / Public Health / School Mental Health journals. '
  'Có thể là target Q2-Q3 IF.')

H3('C5. Em đề xuất Q4 = Phương án A (LPA)')

REC('Em đề xuất Q4 = Phương án A (Latent profile analysis) vì ba lý do.')

B('(1) Phương pháp khác hoàn toàn với Q2/Q3 SEM — không trùng '
  'lặp thống kê, không tự đạo văn.', 0)

B('(2) Dữ liệu hiện có (1.352 HS, 8 thang đo) đủ cho LPA 3-5 class. '
  'Không cần thu thập thêm.', 0)

B('(3) Tạo "hat-trick" báo cáo từ một bộ dữ liệu lớn: SEM tích hợp '
  '(Q2) + SEM bất biến giới (Q3) + LPA phân lớp tiềm ẩn (Q4) — bộ ba '
  'này đáp ứng đầy đủ yêu cầu khoa học cho 3 paper riêng biệt cùng '
  'một cohort.', 0)


# ============================================================
H2('PHẦN D — KIẾN NGHỊ CỤ THỂ')

P('Em đề xuất nhóm cân nhắc bốn việc tiếp theo:')

B('Việc 1: Em sẽ bổ sung Section 4.5 "Gender difference in social '
  'anxiety: Reconciliation with cross-national evidence" vào Q2 v1, '
  'sử dụng 6 luận chứng đã soạn trong PHẦN B. Bổ sung khoảng 600–800 '
  'từ tiếng Anh + tương đương tiếng Việt. Trích dẫn Jefferies & Ungar '
  '2020 sẽ thêm vào References (em đã có PDF gốc tại QT035).', 0)

B('Việc 2: Sau khi thầy + chị Hằng confirm chọn Q3 = Phương án B và '
  'Q4 = Phương án A, em sẽ soạn outline IMRaD cho Q3 và Q4.', 0)

B('Việc 3: Em sẽ cập nhật 4VanDe BLOCKING v4 với (a) cập nhật Q3-6 '
  'đạo đức (NCS đã cam kết), (b) thêm Q3-Q4 title list để chị Hằng '
  'và thầy xác nhận trước khi đề xuất IRB cùng lúc cho 2 hoặc 3 bài.', 0)

B('Việc 4: Em sẽ cập nhật PhuongAnXuLy thành phiên bản v2 với '
  'roadmap 3 bài (Q2 → Q3 → Q4) và timeline cụ thể.', 0)


# ============================================================
H2('THAM KHẢO')

refs = [
    '[1] Jefferies P, Ungar M. (2020). Social anxiety in young people: '
    'A prevalence study in seven countries. PLOS ONE, 15(9), e0239133. '
    'DOI: 10.1371/journal.pone.0239133. PDF: '
    '02_Papers-goc/The-gioi_Au-My-Uc/QT035_Jefferies_SocialAnxiety_'
    '7Countries_2020.pdf — em đã đối chiếu trực tiếp.',

    '[2] Hankin BL, Mermelstein R, Roesch L. (2007). Sex differences in '
    'adolescent depression: Stress exposure and reactivity models. '
    'Child Development, 78(1), 279–295. PMID: 17328705.',

    '[3] McLean CP, Anderson ER. (2009). Brave men and timid women? A '
    'review of the gender differences in fear and anxiety. Clinical '
    'Psychology Review, 29(6), 496–505. PMID: 19541399.',

    '[4] Hill JP, Lynch ME. (1983). The intensification of gender-related '
    'role expectations during early adolescence. In: Brooks-Gunn J, '
    'Petersen AC (eds.), Girls at puberty: Biological and psychosocial '
    'perspectives. New York: Plenum Press, pp. 201–228.',

    '[5] Kessler RC, Berglund P, Demler O, Jin R, Merikangas KR, Walters '
    'EE. (2005). Lifetime prevalence and age-of-onset distributions of '
    'DSM-IV disorders in the National Comorbidity Survey Replication. '
    'Archives of General Psychiatry, 62(6), 593–602. PMID: 15939837. '
    'DOI: 10.1001/archpsyc.62.6.593.',

    '[6] Rose AJ. (2002). Co-rumination in the friendships of girls and '
    'boys. Child Development, 73(6), 1830–1843. PMID: 12487497. DOI: '
    '10.1111/1467-8624.00509.',

    '[7] Stankov L. (2010). Unforgiving Confucian culture: A breeding '
    'ground for high academic achievement, test anxiety and self-doubt? '
    'Learning and Individual Differences, 20(6), 555–563. DOI: '
    '10.1016/j.lindif.2010.05.003.',

    '[8] Small S, Blanc J. (2021). Mental Health During COVID-19: Tam '
    'Giao and Vietnam\'s Response. Frontiers in Psychiatry, 11, 589618. '
    'DOI: 10.3389/fpsyt.2020.589618. PMID: 33536961.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75)
    p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# Footer
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Soạn 08/06/2026 — Trả lời tin nhắn thầy Nguyễn Minh Đức '
              'ngày 08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')

# Quick word count
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            chunks.append(cell.text)
total_words = sum(len(p.split()) for p in chunks)
print(f'WORD COUNT: {total_words}')
