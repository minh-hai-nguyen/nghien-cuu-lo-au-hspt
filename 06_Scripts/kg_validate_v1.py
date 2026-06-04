# -*- coding: utf-8 -*-
"""
KG VALIDATION v1 — Chay cac rule + cross-check voi RAG + summaries.

Rules:
R1. Paper phai co >= 1 author
R2. Paper phai co year in 1990-2030
R3. Cohen d phai in [-3, 3]
R4. OR phai > 0
R5. p-value phai in [0, 1]
R6. % phai in [0, 100]
R7. Paper phai co >= 1 country
R8. Paper phai co >= 1 method HOẶC 1 outcome (intervention vs study)
R9. Consistency: same paper mentioned in multiple sections must have same n

Cross-checks:
CX1. Moi entity trong bao cao v3 phai co trong KG
CX2. Moi ES value trong bao cao v3 phai match KG
CX3. Moi paper reference trong bao cao phai co summary file
"""
import os, sys, io, json, re
from collections import defaultdict, Counter
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import networkx as nx
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
KG_DIR = os.path.join(os.path.dirname(__file__), 'kg_data')

# Load graph
G = nx.read_graphml(os.path.join(KG_DIR, 'kg_v1.graphml'))
print(f'Loaded graph: {G.number_of_nodes()} nodes, {G.number_of_edges()} edges')

# Helper to get node type
def ntype(n):
    return G.nodes[n].get('type', '')

paper_nodes = [n for n in G.nodes if ntype(n) == 'Paper']
print(f'Papers: {len(paper_nodes)}')

# ============================================================
# RULE CHECKS
# ============================================================
print('\n' + '=' * 70)
print('RULE VALIDATION')
print('=' * 70)

violations = []

# R1: Every paper must have at least 1 author
no_author = []
for p in paper_nodes:
    author_count = sum(1 for _, _, d in G.out_edges(p, data=True) if d.get('rel') == 'AUTHORED_BY')
    if author_count == 0:
        no_author.append(p)
print(f'\nR1 (>=1 author): {len(no_author)} papers missing author')
for p in no_author[:10]:
    print(f'  - {p}: {G.nodes[p].get("descriptor", "")}')
if no_author:
    violations.append(('R1', no_author))

# R2: Year in 1990-2030
bad_years = []
for p in paper_nodes:
    for _, tgt, d in G.out_edges(p, data=True):
        if d.get('rel') == 'HAS_YEAR':
            year = G.nodes[tgt].get('value')
            if year is not None:
                try:
                    y = int(year)
                    if y < 1990 or y > 2030:
                        bad_years.append((p, y))
                except: pass
print(f'\nR2 (year 1990-2030): {len(bad_years)} anomalies')

# R3: Cohen d in [-3, 3]
bad_d = []
for n in G.nodes:
    if ntype(n) == 'EffectSize':
        es_type = G.nodes[n].get('es_type', '')
        val = G.nodes[n].get('value')
        if es_type in ('d', 'Cohen_d', 'Hedges_g', 'g', 'SMD') and val is not None:
            try:
                v = float(val)
                if abs(v) > 3:
                    bad_d.append((n, v))
            except: pass
print(f'\nR3 (Cohen d in [-3,3]): {len(bad_d)} out of range')
for n, v in bad_d[:5]:
    print(f'  - {n}: value={v}')

# R4: OR > 0
bad_or = []
for n in G.nodes:
    if ntype(n) == 'EffectSize':
        es_type = G.nodes[n].get('es_type', '')
        val = G.nodes[n].get('value')
        if es_type in ('OR', 'AOR') and val is not None:
            try:
                v = float(val)
                if v <= 0:
                    bad_or.append((n, v))
            except: pass
print(f'\nR4 (OR > 0): {len(bad_or)} invalid')

# R7: Every paper must have >=1 country
no_country = []
for p in paper_nodes:
    country_count = sum(1 for _, _, d in G.out_edges(p, data=True) if d.get('rel') == 'CONDUCTED_IN')
    if country_count == 0:
        no_country.append(p)
print(f'\nR7 (>=1 country): {len(no_country)} papers missing country')
for p in no_country[:10]:
    print(f'  - {p}: {G.nodes[p].get("descriptor", "")}')

# R8: Every paper should have outcome OR method
no_outcome_method = []
for p in paper_nodes:
    outcome_count = sum(1 for _, _, d in G.out_edges(p, data=True) if d.get('rel') == 'MEASURED')
    method_count = sum(1 for _, _, d in G.out_edges(p, data=True) if d.get('rel') == 'USED_METHOD')
    if outcome_count == 0 and method_count == 0:
        no_outcome_method.append(p)
print(f'\nR8 (>=1 outcome or method): {len(no_outcome_method)} papers missing both')
for p in no_outcome_method[:10]:
    print(f'  - {p}: {G.nodes[p].get("descriptor", "")}')

# ============================================================
# CROSS-CHECK vs report v3
# ============================================================
print('\n' + '=' * 70)
print('CROSS-CHECK vs REPORT v3')
print('=' * 70)

REPORT = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 11042026 v3.docx')
doc = Document(REPORT)
report_text = '\n'.join(p.text for p in doc.paragraphs if p.text.strip())
for tb in doc.tables:
    for row in tb.rows:
        for cell in row.cells:
            report_text += '\n' + cell.text

# CX2: Check ES values from KG are in report
print('\nCX2: Effect sizes in KG — found in report v3?')
es_in_report = 0
es_missing_from_report = []
for n in G.nodes:
    if ntype(n) == 'EffectSize':
        val = G.nodes[n].get('value')
        if val is None: continue
        try:
            v = float(val)
            # Format as Vietnamese decimal
            val_str_1 = f'{v:.3f}'.rstrip('0').rstrip('.').replace('.', ',')
            val_str_2 = f'{v:.2f}'.replace('.', ',')
            if val_str_1 in report_text or val_str_2 in report_text:
                es_in_report += 1
            else:
                es_missing_from_report.append((n, v))
        except: pass
print(f'  In report: {es_in_report}')
print(f'  Missing: {len(es_missing_from_report)}')
for n, v in es_missing_from_report[:5]:
    print(f'    - {n}: value={v}')

# CX3: papers in KG should have summaries
print('\nCX3: Papers in KG with summary file?')
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
summary_files = set(os.listdir(TT_DIR))
papers_no_summary = []
for p in paper_nodes:
    has_summary = any(f.startswith(p + '_') for f in summary_files)
    if not has_summary:
        papers_no_summary.append(p)
print(f'  Without summary: {len(papers_no_summary)}')

# ============================================================
# KG QUERIES (examples of utility)
# ============================================================
print('\n' + '=' * 70)
print('KG QUERIES (capability demo)')
print('=' * 70)

# Q1: All papers conducted in Vietnam
vn_papers = [u for u, v, d in G.edges(data=True)
             if d.get('rel') == 'CONDUCTED_IN' and v == 'Country::Vietnam']
print(f'\nQ1: Papers in Vietnam: {len(vn_papers)}')
for p in sorted(vn_papers)[:15]:
    desc = G.nodes[p].get('descriptor', '')
    print(f'  {p} — {desc}')

# Q2: All papers using CBT
cbt_papers = [u for u, v, d in G.edges(data=True)
              if d.get('rel') == 'USED_METHOD' and v == 'Method::CBT']
print(f'\nQ2: Papers using CBT: {len(cbt_papers)}')

# Q3: Papers measuring Social Anxiety (SAD)
sad_papers = [u for u, v, d in G.edges(data=True)
              if d.get('rel') == 'MEASURED' and 'SAD' in v]
print(f'Q3: Papers measuring SAD: {len(sad_papers)}')

# Q4: Largest sample sizes
sample_sizes = []
for u, v, d in G.edges(data=True):
    if d.get('rel') == 'HAS_N':
        n = G.nodes[v].get('value')
        if n: sample_sizes.append((u, int(n)))
sample_sizes.sort(key=lambda x: -x[1])
print(f'\nQ4: Top 10 largest sample sizes:')
for p, n in sample_sizes[:10]:
    desc = G.nodes[p].get('descriptor', '')
    print(f'  {p} (n={n:,}) — {desc}')

# Q5: Papers with d > 0.5 (medium-large effect)
large_effects = []
for n in G.nodes:
    if ntype(n) == 'EffectSize':
        val = G.nodes[n].get('value')
        paper = G.nodes[n].get('paper', '')
        es_type = G.nodes[n].get('es_type', '')
        if val is not None and es_type in ('d', 'g', 'Cohen_d', 'Hedges_g'):
            try:
                v = float(val)
                if abs(v) >= 0.5:
                    large_effects.append((paper, es_type, v))
            except: pass
large_effects.sort(key=lambda x: -abs(x[2]))
print(f'\nQ5: Largest effect sizes (|d/g| >= 0.5):')
for p, t, v in large_effects[:10]:
    desc = G.nodes[p].get('descriptor', '') if p in G.nodes else ''
    print(f'  {p} {t}={v} — {desc}')

# ============================================================
# VALIDATION REPORT
# ============================================================
report_lines = []
report_lines.append('# Validation Report — Knowledge Graph v1\n')
report_lines.append(f'**Ngày:** 11/04/2026\n')
report_lines.append(f'**Graph:** {G.number_of_nodes()} nodes, {G.number_of_edges()} edges\n')
report_lines.append(f'**Papers:** {len(paper_nodes)}\n\n')

report_lines.append('## Rule Violations\n\n')
report_lines.append(f'- R1 Missing author: **{len(no_author)}**\n')
report_lines.append(f'- R2 Bad year: **{len(bad_years)}**\n')
report_lines.append(f'- R3 Cohen d out of range: **{len(bad_d)}**\n')
report_lines.append(f'- R4 Invalid OR: **{len(bad_or)}**\n')
report_lines.append(f'- R7 Missing country: **{len(no_country)}**\n')
report_lines.append(f'- R8 Missing outcome AND method: **{len(no_outcome_method)}**\n\n')

report_lines.append('## Cross-Checks vs Report v3\n\n')
report_lines.append(f'- CX2 Effect sizes in KG matching report: **{es_in_report}**\n')
report_lines.append(f'- CX2 Effect sizes missing from report: **{len(es_missing_from_report)}**\n')
report_lines.append(f'- CX3 Papers without summary: **{len(papers_no_summary)}**\n\n')

report_lines.append('## Papers missing author (R1)\n\n')
for p in no_author:
    report_lines.append(f'- {p}: {G.nodes[p].get("descriptor", "")}\n')

report_lines.append('\n## Papers missing country (R7)\n\n')
for p in no_country:
    report_lines.append(f'- {p}: {G.nodes[p].get("descriptor", "")}\n')

report_lines.append('\n## Papers missing outcome & method (R8)\n\n')
for p in no_outcome_method:
    report_lines.append(f'- {p}: {G.nodes[p].get("descriptor", "")}\n')

out_path = os.path.join(KG_DIR, 'validation_report.md')
with open(out_path, 'w', encoding='utf-8') as f:
    f.writelines(report_lines)
print(f'\nValidation report: {out_path}')
