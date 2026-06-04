# -*- coding: utf-8 -*-
"""Bài 8: Wen et al. 2020 — CTH v5 (rút gọn ≤2,5 trang)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 8 =====
bl('Tóm tắt bài 8', bold=True)

rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình \u00ab Phân tích hồ sơ tiềm ẩn về lo âu ở học sinh trung học cơ sở tại các vùng nông thôn kém phát triển của Trung Quốc \u00bb (A Latent Profile Analysis of Anxiety among Junior High School Students in Less Developed Rural Regions of China), do Xiaotong Wen và cộng sự (2020), thuộc Đại học Nam Xương, tỉnh Giang Tây, khảo sát trên 900 học sinh lớp 9\u201312 (tuổi TB 14,14) tại 6 huyện nông thôn tỉnh Giang Tây, miền Đông Trung Quốc, xuất bản trên Int J Environ Res Public Health, 17(11):4079.')

rb('Phương pháp nghiên cứu')
bl('Công trình sử dụng Trắc nghiệm Sức khỏe Tâm thần (MHT \u2014 Mental Health Test, 100 mục, 8 thang nội dung) \u2014 thang chẩn đoán lo âu tiêu chuẩn hóa dành cho học sinh tiểu học và THCS tại Trung Quốc (Cronbach alpha 0,878). Nói cách khác, nghiên cứu sử dụng phân tích hồ sơ tiềm ẩn (LPA) kết hợp hồi quy logistic đa biến \u2014 phương pháp tập trung vào cá nhân (person-centered) thay vì biến (variable-centered).', bold=True)
bl('Lấy mẫu cụm ngẫu nhiên phân tầng nhiều giai đoạn: Ủy ban Y tế tỉnh chọn huyện → chọn xã → chọn trường → chọn ngẫu nhiên lớp. Phân tích bằng Mplus 7.4 (LPA) và SPSS 24.0 (hồi quy logistic).')

rb('Kết quả nghiên cứu định lượng')
bl('LPA xác định 3 hồ sơ lo âu (N = 900):', bold=True)
tbl(
    ['Hồ sơ', 'n', 'Tỷ lệ', 'Đặc điểm'],
    [
        ['Lo âu nhẹ', '173', '19,2%', 'Điểm thấp nhất trên 8 yếu tố'],
        ['Lo âu trung bình', '504', '56,0%', 'Điểm trung bình'],
        ['Lo âu nặng', '223', '24,8%', 'Điểm cao nhất trên 8 yếu tố'],
    ],
    widths=[3.5, 1.5, 2.0, 5.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan (hồi quy logistic, nhóm tham chiếu: lo âu nhẹ):', bold=True)
tbl(
    ['Yếu tố', 'Lo âu TB (OR)', 'Lo âu nặng (OR)'],
    [
        ['Áp lực học tập rất cao', '6,523*', '11,579*'],
        ['Áp lực học tập cao', '6,122*', '5,894*'],
        ['Giới tính nam (vs nữ)', '0,649* (bảo vệ)', '0,262* (bảo vệ)'],
        ['SKTT trường đầy đủ', '\u2014', '0,562* (bảo vệ)'],
        ['Thành tích xuất sắc', '0,377* (bảo vệ)', '\u2014'],
    ],
    widths=[4.5, 3.5, 3.5]
)
doc.add_paragraph()

bh3('Khác biệt giới tính : Trái ngược với đa số y văn (nữ > nam), nghiên cứu này cho thấy nam có nguy cơ lo âu nặng chỉ bằng 0,262 lần so với nữ (P < 0,05) \u2014 tức nữ lo âu nặng gấp gần 4 lần nam. Phát hiện này phù hợp với xu hướng quốc tế nhưng mức chênh lệch lớn bất thường.')

rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bl('*Áp lực học tập \u2014 yếu tố nguy cơ mạnh nhất.* Học sinh tự đánh giá áp lực "rất cao" có nguy cơ lo âu nặng gấp 11,6 lần. Đây là OR lớn nhất trong tất cả 11 bài nghiên cứu, phản ánh đặc thù hệ thống giáo dục Trung Quốc tập trung vào thi cử.')
bl('*Hỗ trợ SKTT tại trường \u2014 yếu tố bảo vệ.* Khi trường có dịch vụ SKTT đầy đủ, nguy cơ lo âu nặng giảm (OR = 0,562), gợi ý can thiệp hiệu quả dựa trên trường học.')

rb('Kết luận')
bl('Dữ liệu cho thấy 24,8% học sinh THCS nông thôn Trung Quốc có lo âu nặng, với áp lực học tập là yếu tố nguy cơ mạnh nhất (OR = 11,6). Hỗ trợ SKTT tại trường giảm nguy cơ lo âu nặng (OR = 0,562). Nữ sinh có nguy cơ cao gấp gần 4 lần nam sinh. Giảm áp lực học tập và thiết lập phòng tư vấn tâm lý chuyên nghiệp tại trường là hai khuyến nghị chính.', bold=True)

rh2('Phản biện')
bl('Điểm mạnh: phương pháp LPA độc đáo, tập trung vào cá nhân thay vì biến. Tuy nhiên, cỡ mẫu 900 khiêm tốn, chỉ khảo sát nông thôn một tỉnh. MHT là thang đo Trung Quốc, khó so sánh quốc tế (không dùng GAD-7/DASS-21). Tạp chí IJERPH (PubMed-indexed) nhưng impact factor trung bình.')

rh2('Hướng nghiên cứu tiếp theo')
bl('Nghiên cứu dọc theo dõi lo âu từ THCS đến THPT. So sánh nông thôn\u2013thành thị cùng tỉnh. Can thiệp giảm áp lực học tập và đánh giá hiệu quả.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50 Trung bình-Khá. LPA độc đáo nhưng cỡ mẫu nhỏ, công cụ đặc thù Trung Quốc.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('08_Wen_et_al_2020.docx')
print('08_Wen_et_al_2020.docx OK')
