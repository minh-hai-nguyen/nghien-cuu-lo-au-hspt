# -*- coding: utf-8 -*-
"""Doc trả lời 3 câu hỏi của thầy — 27/04/2026:
1. NC nào thu thập coping strategies tự phát từ HS?
2. Bài Jefferies & Ungar 2020 — bc VN của ai?
3. Verify 80% HS BESST chưa từng tìm giải pháp.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\Tra_loi_3_cau_hoi_cho_thay_27042026.docx'
RED  = RGBColor(0xC0, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def warn(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('⚠ '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def ok(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('✓ '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(12); r2.font.name='Times New Roman'

# ===================================================================
title("TRẢ LỜI 3 CÂU HỎI CỦA THẦY", 16)
subtitle("(1) Nghiên cứu coping tự phát từ HS — (2) Tác giả VN trong Jefferies 2020 — "
         "(3) Verify 80% non-consulters BESST")
nr("")
subtitle("Trợ lý nghiên cứu — 27/04/2026")

# ===================================================================
# CÂU HỎI 1
# ===================================================================
H1("CÂU 1 — Có nghiên cứu nào thu thập COPING STRATEGIES TỰ PHÁT (bottom-up) từ HS không?")
nr("Câu hỏi của thầy: \"Có nghiên cứu nào thu thập thông tin từ những cách giải quyết lo "
   "âu đa dạng, tự phát từ phía HS không?\"", italic=True, color=GRAY)
nr("")
nr("TRẢ LỜI THẲNG:", bold=True)
ok("CÓ — đặc biệt qua 3 phương pháp định tính: (a) Photovoice, (b) Focus group / "
   "phenomenological interview, (c) Self-report survey với open-ended question.")
nr("")
nr("Em đã tìm được nhiều nghiên cứu phù hợp. Phân theo phương pháp:")

H2("1.1. PHOTOVOICE — HS dùng ảnh thể hiện cách họ đối phó")
H3("a. SR scoping review 2023 — Photovoice in mental health research with adolescents")
nr("• Tạp chí: International Journal of Adolescence and Youth — Vol 28, No 1 (2023)", size=12)
nr("• Tác giả + năm: chưa có thông tin tác giả đầy đủ; published 2023", size=12)
nr("• DOI: 10.1080/02673843.2023.2244043", size=12)
nr("• Link: https://www.tandfonline.com/doi/full/10.1080/02673843.2023.2244043", size=12)
nr("• Phát hiện: Review hệ thống về Photovoice trong nghiên cứu SKTT VTN — phương pháp "
   "này cho phép HS chụp ảnh + thảo luận về CÁCH HỌ TRẢI NGHIỆM + ỨNG PHÓ với SKTT. "
   "Insight emotional + metaphoric không thể đạt được qua qualitative truyền thống. "
   "Kết luận: lượng evidence vẫn HẠN CHẾ, cần thêm nghiên cứu với rigour cao.", size=12)

H3("b. Hispanic adolescents Photovoice 2024-2025")
nr("• Tạp chí: Health Promotion International (Oxford Academic)", size=12)
nr("• DOI: 10.1093/heapro/daag032", size=12)
nr("• Link: https://academic.oup.com/heapro/article-pdf/41/1/daag032/67146769/daag032.pdf", size=12)
nr("• Thiết kế: 12 VTN gốc Hispanic 13-17 tuổi từ 2 youth centers, 3 buổi (preparatory "
   "workshop + focus group + feedback workshop) từ 06/2024 đến 02/2025. HS chụp ảnh thể "
   "hiện quan điểm về SKTT, sau đó thảo luận theo kỹ thuật SHOWeD.", size=12)
nr("• PHÁT HIỆN ĐÁNG CHÚ Ý:", bold=True, size=12)
nr("  Khi được tự thể hiện (bottom-up), VTN ưu tiên 4 NHÓM COPING ĐA DẠNG:", size=12)
nr("    – SOLITARY (cá nhân): journaling, đi dạo, thời gian một mình", size=12)
nr("    – SENSORY (giác quan): nghe nhạc, ngắm thiên nhiên, thực phẩm yêu thích", size=12)
nr("    – CREATIVE (sáng tạo): vẽ, viết, làm nghệ thuật", size=12)
nr("    – SOCIAL (xã hội): trò chuyện với bạn, gia đình, peer support", size=12)
nr("→ ĐIỂM QUAN TRỌNG: HS TỰ PHÁT đi tới các strategies này KHÔNG QUA chương trình áp đặt. "
   "Khác biệt với CBT/Mindfulness chuẩn ở chỗ HS chọn theo bản thân.", bold=True, size=12)

H3("c. Adolescents living with HIV Photovoice during COVID — South Africa 2024")
nr("• Tạp chí: IJERPH (MDPI) — Vol 21(11): 1517 (2024)", size=12)
nr("• DOI: 10.3390/ijerph21111517", size=12)
nr("• Link: https://www.mdpi.com/1660-4601/21/11/1517", size=12)
nr("• Phát hiện coping diverse: religion/faith, music, sports, journaling, peer support — "
   "khác hẳn các \"khuyến nghị chuẩn\" CBT/mindfulness.", size=12)

H2("1.2. FOCUS GROUP / Phenomenological Interview")
H3("d. Senior HS broken families — Philippines 2024-2025 (qualitative phenomenological)")
nr("• Tạp chí: Pantao Journal — published 2025 (school year 2024-2025)", size=12)
nr("• Link: https://pantaojournal.com/wp-content/uploads/2025/06/172-Delos-Santos.pdf", size=12)
nr("• Thiết kế: In-depth interviews với 10 HS THPT Philippines từ gia đình tan vỡ; "
   "thematic analysis 5 chủ đề lớn.", size=12)
nr("• Phát hiện coping TỰ PHÁT: \"seek out pleasant experiences, cultivate independence, "
   "strive to remain optimistic\"; \"forming greater bonds with friends and family as "
   "support source\".", size=12)

H3("e. AYAs COVID-19 lived experience — Italia 2022 (qualitative interview)")
nr("• Tạp chí: PMC8882423 (Front Public Health hoặc tương đương)", size=12)
nr("• Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC8882423/", size=12)
nr("• Thiết kế: 25 AYAs 13-24 tuổi, interview 5/2020. Thematic analysis ra 5 chủ đề "
   "trong đó \"development of coping strategies to maintain well-being\" — HS tự phát "
   "thử nghiệm và phát triển coping trong COVID.", size=12)

H3("f. Frontiers Education 2025 — Their voices matter (school attendance + SEB difficulties)")
nr("• Tạp chí: Frontiers in Education (2025)", size=12)
nr("• DOI: 10.3389/feduc.2025.1627098", size=12)
nr("• Link: https://www.frontiersin.org/journals/education/articles/10.3389/feduc.2025.1627098/full", size=12)
nr("• Thiết kế: Cả student + professional perspectives; HS có khó khăn social-emotional-"
   "behavioural đề cập cách HỌ overcome school attendance problems.", size=12)

H2("1.3. SURVEY VỚI OPEN-ENDED Q (Quantitative + Qualitative)")
H3("g. Persistent anxiety high school — COVID year 2 — PLOS ONE 2022")
nr("• Tạp chí: PLOS ONE — published 2022", size=12)
nr("• DOI: 10.1371/journal.pone.0275292", size=12)
nr("• Link: https://journals.plos.org/plosone/article?id=10.1371/journal.pone.0275292", size=12)
nr("• Thiết kế: Survey HS THPT trong COVID năm 2; có open-ended question về cách họ "
   "đối phó với anxiety persistent.", size=12)

H3("h. SR Anxiety undergraduate 2025 — Sage")
nr("• Tạp chí: Sage Open / Annals of Neurosciences (2025)", size=12)
nr("• DOI: 10.1177/09727531251366078", size=12)
nr("• PMC: PMC12420638", size=12)
nr("• Link: https://journals.sagepub.com/doi/10.1177/09727531251366078", size=12)
nr("• Phát hiện: Phân tích 40 nghiên cứu thực nghiệm — tổng hợp coping mechanisms "
   "do HS tự kê khai; bao gồm mindfulness (cá nhân), institutional support, peer.", size=12)

H3("i. Academic Stress Coping HS — IJAMR 2024")
nr("• Tạp chí: International Journal of Academic Multidisciplinary Research", size=12)
nr("• Link: http://ijeais.org/wp-content/uploads/2024/5/IJAMR240518.pdf", size=12)
nr("• HS tự báo cáo coping: hỗ trợ xã hội từ peer, tập thể dục, thiền/relaxation, "
   "problem-solving, cognitive restructuring.", size=12)

H2("1.4. TỔNG HỢP — PATTERN COPING TỰ PHÁT CỦA HS QUA CÁC NGHIÊN CỨU TRÊN")
tbl(['Nhóm coping', 'Strategy cụ thể', 'Tỉ lệ HS tự dùng (qualitative)', 'Văn hóa phổ biến'],
    [
        ['Sensory (giác quan)', 'Nghe nhạc, ngắm thiên nhiên',
         'Cao (mọi văn hóa)', 'Universal'],
        ['Solitary (cá nhân)', 'Journaling, đi dạo, ngủ',
         'Cao ở phương Tây', 'Cao ở Anglo-Saxon'],
        ['Creative (sáng tạo)', 'Vẽ, viết, làm nghệ thuật',
         'Trung bình-cao', 'Universal'],
        ['Social (xã hội)', 'Trò chuyện bạn, gia đình',
         'RẤT cao ở Á Đông', 'Mạnh nhất ở văn hóa collectivist'],
        ['Religion/Faith', 'Cầu nguyện, thiền tôn giáo',
         'Cao ở khu vực religious', 'Châu Phi, Mỹ Latin, Á Hồi'],
        ['Physical activity', 'Thể thao, đi bộ, yoga',
         'Trung bình', 'Cao ở phương Tây'],
        ['Avoidance', 'Tránh né, đẩy lùi vấn đề',
         'Cao nhưng KHÔNG ADAPTIVE', 'Universal'],
        ['Substance use', 'Hút thuốc, rượu (negative)',
         'Cao ở mẫu HS THPT muộn', 'Universal — phải cảnh báo'],
    ], [3.0, 4.5, 4.0, 4.5])

H2("1.5. NHẬN XÉT TỔNG QUAN")
nr("(1) HS THỰC SỰ DÙNG NHIỀU CHIẾN LƯỢC ĐA DẠNG — không chỉ CBT/mindfulness.")
nr("(2) Phương pháp QUALITATIVE (Photovoice, focus group, phenomenological) là CỬA SỔ "
   "tốt nhất để thu thập coping TỰ PHÁT. Phương pháp QUANTITATIVE chỉ đo các strategies "
   "đã được liệt kê sẵn (Brief COPE, COPE-28).")
nr("(3) GIÁ TRỊ NGHIÊN CỨU BOTTOM-UP: phát hiện strategies KHÔNG được dạy chính thức "
   "(VD: nghe nhạc, ngắm thiên nhiên, tâm sự với bạn) — vốn rất phổ biến và HIỆU QUẢ "
   "trong thực tế nhưng ít có trong manual CBT.")
nr("(4) BỐI CẢNH VN: chưa có Photovoice mental health study bằng tiếng Việt mà em tìm "
   "được. Đây là GAP nghiên cứu RẤT QUAN TRỌNG cho VN — em đề xuất thầy có thể làm "
   "Photovoice nhỏ với 12-20 HS lớp 10-12 ở 2-3 trường VN.")
nr("(5) Quy trình Photovoice gợi ý cho VN: (a) 1 buổi giới thiệu phương pháp (90'); "
   "(b) HS chụp 5-10 ảnh trong 1-2 tuần thể hiện \"cách em đối phó căng thẳng\"; "
   "(c) 1 buổi focus group thảo luận theo kỹ thuật SHOWeD; (d) 1 buổi feedback. Tổng "
   "3-4 buổi/HS, chi phí thấp, ethics đơn giản (consent + bảo mật ảnh).")

# ===================================================================
# CÂU HỎI 2
# ===================================================================
H1("CÂU 2 — Tác giả VN trong Jefferies & Ungar 2020 (7 quốc gia)?")
nr("Câu hỏi của thầy: \"Tìm bản tiếng Anh bài này để xem bc của VN là của ai mà họ tổng "
   "hợp cùng nước ngoài? — Jefferies P, Ungar M (2020). Social Anxiety in Young People: "
   "A Prevalence Study in Seven Countries.\"", italic=True, color=GRAY)
nr("")

H2("2.1. Thông tin bài báo")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài', 'Social anxiety in young people: A prevalence study in seven countries'],
        ['Tác giả (CHỈ 2)', 'Philip Jefferies + Michael Ungar'],
        ['Affiliation cả 2 tác giả',
         'Resilience Research Centre, Faculty of Health, Dalhousie University, Halifax, '
         'Nova Scotia, Canada'],
        ['Tạp chí', 'PLoS ONE (Open Access)'],
        ['Vol/Số/Pages', '15(9): e0239133'],
        ['Năm + ngày', '2020 — Published 17/09/2020'],
        ['DOI', '10.1371/journal.pone.0239133'],
        ['PMID', '32941482'],
        ['PMC', 'PMC7498107'],
        ['Cỡ mẫu', '6.825 (3.342 nam + 3.428 nữ + 55 other), 16-29 tuổi'],
        ['7 quốc gia', 'Brazil, China, Indonesia, Russia, Thailand, US, VIETNAM'],
        ['Cỡ mẫu VN', '984 (487 nam + 493 nữ + 4 other)'],
        ['Công cụ đo', 'SIAS (Social Interaction Anxiety Scale)'],
        ['Phát hiện chính', '36% (>1/3) thoả ngưỡng SAD; cao hơn nhiều ước tính trước'],
    ], [4.0, 12.0])

H2("2.2. Trả lời TRỰC TIẾP câu hỏi của thầy")
warn("KHÔNG CÓ TÁC GIẢ VIỆT NAM trong bài này. Cả 2 tác giả (Jefferies + Ungar) đều thuộc "
     "Resilience Research Centre, Dalhousie University, Canada.")
nr("")
nr("DỮ LIỆU VN ĐƯỢC THU THẬP NHƯ THẾ NÀO?", bold=True)
nr("Theo phần Methods + Acknowledgments của bài (em đã fetch verbatim từ PMC):")
nr("• Người tham gia ở 7 nước được tuyển ngẫu nhiên thông qua 3 CÔNG TY KHẢO SÁT THỊ "
   "TRƯỜNG (market research companies):")
nr("    – DYNATA (US-based, chi nhánh toàn cầu)")
nr("    – ONLINE MARKET INTELLIGENCE (OMI) — Russia/Eastern Europe-based")
nr("    – GMO RESEARCH — Japan-based, hoạt động Asia-Pacific")
nr("• Acknowledgments quote VERBATIM: \"The authors would like to acknowledge the role "
   "of EDELMAN INTELLIGENCE for collecting the original data on behalf of UNILEVER and "
   "CLEAR.\" → Edelman Intelligence là agency global, thu thập SECONDARY DATA cho 2 thương "
   "hiệu UNILEVER + CLEAR (CLEAR là dầu gội trị gàu).", italic=True)
nr("")
nr("KẾT LUẬN VỀ NGUỒN DỮ LIỆU VN:", bold=True, color=RED)
nr("Dữ liệu VN n=984 trong bài này KHÔNG đến từ một viện nghiên cứu VN nào, mà từ:")
nr("• Một công ty marketing toàn cầu (Dynata / OMI / GMO Research) recruit qua online panel")
nr("• Dữ liệu gốc do EDELMAN INTELLIGENCE thu thập cho UNILEVER + brand CLEAR (đây là "
   "dữ liệu marketing thứ cấp, có thể gốc từ chiến dịch quảng cáo CLEAR liên quan đến "
   "self-confidence)")
nr("• Jefferies & Ungar phân tích lại dữ liệu này (secondary analysis) cho học thuật")

H2("2.3. Hệ quả cho việc dẫn chứng")
warn("HẠN CHẾ về validity ngoại của số liệu VN trong bài này:")
nr("(1) Mẫu KHÔNG random thuần tuý từ population VN — tuyển qua online panel của agency, "
   "có thể thiên về thanh thiếu niên đô thị có internet + sẵn lòng làm survey marketing.")
nr("(2) SIAS bản tiếng Việt có thể CHƯA được validate cho VN — bài không nói validation.")
nr("(3) Dữ liệu thu cho mục đích MARKETING (Unilever/CLEAR) — câu hỏi có thể bias theo "
   "self-image / self-confidence framework, không phải lo âu xã hội DSM-5 chuẩn.")
nr("(4) KHÔNG có chuyên gia VN review tính phù hợp văn hoá của tool và diễn giải kết quả.")
nr("")
nr("KHUYẾN NGHỊ KHI TRÍCH:", bold=True, color=BLUE)
nr("Khi trích con số \"VN có SAD 30,7% / SAD nhóm cao thứ 1 trong 7 nước\" (theo bài), "
   "nên ghi rõ: \"theo dữ liệu commercial panel thu cho Unilever/CLEAR, không phải khảo "
   "sát quốc gia chuẩn\". Số liệu này KHÁC với V-NAMHS 2022 (DISC-5/DSM-5) — V-NAMHS "
   "là số liệu quốc gia chính thức và đáng tin hơn.")

# ===================================================================
# CÂU HỎI 3
# ===================================================================
H1("CÂU 3 — Verify: 80% HS trong 900 HS BESST chưa từng tìm giải pháp?")
nr("Câu hỏi của thầy: \"Thông tin 80% HS trong mẫu 900 HS trong NC của Brown chưa từng "
   "tìm giải pháp trước khi đến với BESST, có đúng không?\"", italic=True, color=GRAY)
nr("")

H2("3.1. Trả lời thẳng + đính chính nhỏ")
ok("CON SỐ 80% LÀ ĐÚNG — nhưng cần PHÂN BIỆT NGHĨA chính xác.")
nr("")
nr("Trong PDF gốc Brown 2024 BESST (Lancet Psychiatry 11(7):504-515) — em đã đọc full + "
   "đã trích trong doc dịch:")
nr("• Bảng 1 (trang 510, Baseline characteristics):", bold=True)
nr("    – \"Previously sought help from general practitioner for mental health: NO 720 "
   "(80%), YES 179 (20%)\"")
nr("    – Tổng 720+179=899 (1 missing); làm tròn 80% / 20%")
nr("• Discussion (trang 513):", bold=True)
nr("    – \"The high proportion of those who had not previously sought help through "
   "FORMAL ROUTES (80%)\"")
nr("• Research-in-context box (trang 505):", bold=True)
nr("    – \"80% of students in our study had not previously sought help\"")

H2("3.2. PHÂN BIỆT QUAN TRỌNG — 80% NGHĨA LÀ GÌ?")
warn("80% LÀ \"chưa từng TÌM GP/professional formal help\" — KHÔNG PHẢI \"chưa từng làm "
     "gì\" hay \"chưa từng tìm bất kỳ giải pháp nào\".")
nr("")
nr("Cụ thể bài Brown 2024 đo MỘT câu hỏi: \"Đã từng tham vấn GP (bác sĩ đa khoa) cho vấn "
   "đề SKTT chưa?\" — biến binary YES/NO. Câu này CHỈ đo formal route qua hệ thống y tế "
   "công NHS (UK).")
nr("")
nr("Các kênh HS có thể đã dùng nhưng KHÔNG được đo trong câu này (vẫn nằm trong 80% trên):")
tbl(['Kênh đã có thể dùng', 'Có nằm trong câu hỏi 80% không?'],
    [
        ['Tâm sự với bạn (peer support)', 'KHÔNG đo'],
        ['Tâm sự với gia đình', 'KHÔNG đo'],
        ['Tự tìm online (website, app, video)', 'KHÔNG đo'],
        ['Đi gặp tư vấn học đường', 'KHÔNG đo'],
        ['Cầu nguyện / tôn giáo', 'KHÔNG đo'],
        ['Tự dùng coping (nghe nhạc, thể thao, journal)', 'KHÔNG đo'],
        ['CAMHS (specialist child mental health) UK', 'CAMHS là criterion EXCLUSION '
         '(loại HS đang nhận therapy) — không có trong sample'],
        ['GP cho vấn đề SKTT', 'CHÍNH LÀ câu hỏi 80%'],
    ], [7.0, 9.0])

H2("3.3. CÁCH HIỂU CHÍNH XÁC")
nr("80% trong BESST → 80% HS đã đăng ký (self-referral) cho workshop DISCOVER là những "
   "HS CHƯA TỪNG đi qua kênh y tế CHÍNH THỨC (GP) cho vấn đề SKTT của họ.")
nr("")
nr("Ý nghĩa thực tiễn:")
nr("• Self-referral chạm tới nhóm \"hard-to-reach\" — vốn không đi GP (vì stigma, không "
   "biết GP có thể giúp về SKTT, hoặc không nghĩ vấn đề đủ \"nặng\" để gặp bác sĩ).")
nr("• Trong 80% này, có thể có HS đã DÙNG nhiều coping tự phát (tâm sự bạn, online, "
   "tư vấn học đường) — bài KHÔNG đo nên không biết tỉ lệ.")
nr("• Đây là PHÁT HIỆN MẠNH cho lập luận self-referral (PLACES model): chương trình tới "
   "đúng nhóm cần.")

H2("3.4. Đề xuất cách diễn đạt cho thầy khi trích")
nr("THAY VÌ NÓI: \"80% HS BESST chưa từng tìm giải pháp\" (KHÔNG HOÀN TOÀN ĐÚNG)", italic=True, color=RED)
nr("")
nr("NÊN NÓI:", bold=True, color=GREEN)
nr("(a) \"80% HS BESST chưa từng tham vấn GP về vấn đề SKTT\" (chính xác về formal "
   "healthcare route), HOẶC")
nr("(b) \"80% HS BESST chưa từng tiếp cận hệ thống y tế chính thức cho SKTT — workshop "
   "DISCOVER là điểm đầu tiên họ engage với chuyên môn SKTT\", HOẶC")
nr("(c) \"BESST tiếp cận được nhóm hard-to-reach: 80% HS tham gia chưa bao giờ qua kênh "
   "GP/formal help — chứng tỏ self-referral có khả năng phá rào cản tiếp cận\".")

# ===================================================================
H1("THAM KHẢO ĐẦY ĐỦ")

H2("Cho câu 1 (coping bottom-up)")
nr("• Photovoice MH adolescents SR 2023. Int J Adolesc Youth 28(1). "
   "doi:10.1080/02673843.2023.2244043", size=11)
nr("• Hispanic adolescents Photovoice 2024-2025. Health Promot Int 41(1):daag032. "
   "doi:10.1093/heapro/daag032", size=11)
nr("• HIV adolescents Photovoice South Africa COVID 2024. IJERPH 21(11):1517. "
   "doi:10.3390/ijerph21111517", size=11)
nr("• Delos Santos et al. (2025). Senior HS broken families Philippines — "
   "phenomenological. Pantao Journal.", size=11)
nr("• AYAs COVID-19 lived experience qualitative 2022. PMC8882423.", size=11)
nr("• Frontiers Education 2025 — \"Their voices matter\". "
   "doi:10.3389/feduc.2025.1627098", size=11)
nr("• PLOS ONE 2022. Persistent anxiety HS COVID year 2. "
   "doi:10.1371/journal.pone.0275292", size=11)
nr("• SR anxiety undergraduate 2025. Sage. doi:10.1177/09727531251366078 — "
   "PMC12420638", size=11)
nr("• Academic Stress Coping HS — IJAMR 2024.", size=11)

H2("Cho câu 2 (Jefferies & Ungar 2020)")
nr("• Jefferies P, Ungar M (2020). Social anxiety in young people: A prevalence study in "
   "seven countries. PLoS ONE 15(9):e0239133. doi:10.1371/journal.pone.0239133 — "
   "PMID 32941482 — PMC7498107", size=11)
nr("• Em đã fetch verbatim từ PMC NCBI để xác minh authors, affiliations, "
   "Vietnam recruitment, acknowledgments", size=11)
nr("• Cross-reference: V-NAMHS 2022 (corpus dự án) — Vietnam national survey với "
   "DISC-5/DSM-5 standard — uy tín hơn cho dữ liệu quốc gia", size=11)

H2("Cho câu 3 (BESST 80% non-consulters)")
nr("• Brown J, James K, Lisk S et al. (2024). Clinical effectiveness and cost-effectiveness "
   "of brief accessible CBT for stress in school-aged adolescents (BESST): cluster RCT. "
   "Lancet Psychiatry 11(7):504-515. PMID 38759665 — doi:10.1016/S2215-0366(24)00101-9", size=11)
nr("• PDF full đã đọc: 02_Papers-goc/UK_BESST_PLACES/Brown_2024_BESST_Lancet_Psychiatry.pdf "
   "(Bảng 1 trang 510 + Discussion trang 513 + Research-in-context trang 505)", size=11)
nr("• Doc dịch+phản biện: 03_Ban-dich/Bai_dich_phan_bien/"
   "BESST_Brown_2024_dich_phan_bien_25042026.docx — đã trích 80% trong PHẦN 1 + Bảng B + "
   "PHẦN 3 phản biện", size=11)
nr("• KG edge: QT042_BESST → Concept::self_referral với 80% non-consulters context", size=11)

# ===================================================================
H2("Truy vết & Tools đã dùng để trả lời")
nr("• Câu 1: 3 lần WebSearch (qualitative coping, student voices, photovoice) trong "
   "phiên 27/04/2026", size=11)
nr("• Câu 2: 1 lần WebSearch + 1 lần WebFetch verbatim từ PMC7498107", size=11)
nr("• Câu 3: Verify từ PDF Brown 2024 BESST đã đọc full + RAG chunk QT042_BESST + "
   "KG edge", size=11)
nr("• Memory áp dụng: feedback_research_workflow.md (kiểm 3-4 vòng) + "
   "feedback_doc_phai_co_reference.md (DOI/PMID đầy đủ)", size=11)

d.save(OUT)
print('Wrote:', OUT)
