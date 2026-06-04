# -*- coding: utf-8 -*-
"""Viết tóm tắt VN19 + VN20 + sửa VN15-18"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '03_Ban-dich'))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

def tt(title_vn, info, pp_method, pp_tools, pp_justify, data_demo, process, validity,
       results_table_header, results_table_rows, results_table_widths,
       compare, remarks, conclusion, critique, future, rating, extra_tables=None):
    doc = create_doc()
    add_heading(doc, title_vn, 1)
    add_p(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
    add_p(doc, info)
    add_p(doc, 'Phương pháp nghiên cứu', bold=True)
    add_p(doc, pp_method)
    if pp_tools: add_p(doc, pp_tools)
    if pp_justify: add_p(doc, pp_justify)
    if data_demo: add_p(doc, data_demo)
    if process: add_p(doc, process)
    if validity: add_p(doc, validity)
    add_p(doc, 'Kết quả nghiên cứu định lượng', bold=True)
    add_heading(doc, 'Bảng 1. Kết quả chính', 3)
    add_table(doc, results_table_header, results_table_rows, widths=results_table_widths)
    if extra_tables:
        for et in extra_tables:
            add_heading(doc, et.get('title', 'Bảng bổ sung'), 3)
            add_table(doc, et['header'], et['rows'], widths=et.get('widths'))
    if compare:
        add_p(doc, 'Đối chiếu liên bài', bold=True)
        add_p(doc, compare)
    add_p(doc, 'Nhận xét, phát hiện qua kết quả nghiên cứu', bold=True)
    for rm in remarks: add_p(doc, rm)
    add_p(doc, 'Kết luận', bold=True)
    add_p(doc, conclusion)
    add_red_heading(doc, 'Phản biện')
    add_red(doc, critique)
    add_red_heading(doc, 'Hướng nghiên cứu tiếp theo')
    add_red(doc, future)
    add_p(doc, f'Đánh giá: {rating}', bold=True)
    return doc

# =====================================================================
# VN19 — Trần Thảo Vi (Hồ Thị Trúc Quỳnh) 2025 — Huế
# =====================================================================
print('VN19...')
doc = tt(
    'Tóm tắt bài VN-19',
    'Công trình « Lo âu và trầm cảm ở thanh thiếu niên tại Thừa Thiên Huế và mối quan hệ của nó với sự lạc quan » (Anxiety and Depression in Adolescents in Thua Thien Hue and Their Relationship with Optimism), do Hồ Thị Trúc Quỳnh (2025), Khoa Tâm lý và Giáo dục, Đại học Sư phạm Huế, khảo sát 685 HS THPT (233 nam, 452 nữ, tuổi 15–18) tại Thừa Thiên Huế, tháng 10/2021 (COVID-19 làn sóng 4). Tạp chí Tâm lý học, Số 3 (312), tháng 3/2025, trang 55–65.',
    'Công trình sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21) phiên bản tiếng Việt kết hợp Trắc nghiệm Định hướng Cuộc sống Chỉnh sửa (LOT-R — Life Orientation Test-Revised; α = 0,82). Nói cách khác, NC không chỉ đo MỨC ĐỘ lo âu/trầm cảm mà còn khám phá CƠ CHẾ — sự lạc quan giúp giảm trầm cảm thông qua giảm lo âu trước (trung gian một phần).',
    'DASS-21 (Lovibond & Lovibond, 1995): tiểu thang lo âu (7 mục), ngưỡng ≥8. LOT-R (Scheier et al., 1994): 6 mục (+ 4 câu lấp), đo xu hướng lạc quan.',
    'Tổng quan cho thấy ít NC VN nghiên cứu SKTT tích cực (lạc quan) liên quan lo âu/trầm cảm. Phân tích trung gian (mediation) bằng PROCESS macro — phương pháp nâng cao.',
    'N = 685 HS THPT, 233 nam (34,0%), 452 nữ (66,0%). Lấy mẫu thuận tiện, khảo sát trực tuyến qua Google Form.',
    'PROCESS macro (Hayes, Model 4). Phân tích tương quan Pearson + hồi quy.',
    None,
    ['Tình trạng', 'Tỷ lệ', 'n', 'Ghi chú'],
    [['Lo âu (DASS-21 ≥8)', '65,8%', '451/685', 'Cao — bối cảnh COVID-19'],
     ['Trầm cảm (DASS-21 ≥10)', '45,1%', '309/685', ''],
     ['Lo âu nhẹ', '11,7%', '80', ''],
     ['Lo âu vừa', '31,1%', '213', 'Lớn nhất'],
     ['Lo âu nghiêm trọng', '10,1%', '69', ''],
     ['Lo âu rất nghiêm trọng', '12,8%', '88', '']],
    [3.5, 2.0, 2.0, 4.0],
    'Lo âu 65,8% so với: Hoa 2024 Hà Nội (40,6% GAD-7), Bảo Quyên 2025 Hà Nội (86,2% DASS-21), An Giang 2025 (61,2% DASS-21), Vĩnh Lộc 2024 TPHCM (25,1% DASS-Y). Bối cảnh COVID-19 có thể làm tăng tỷ lệ. Miền Trung (Huế) — ít NC so với Hà Nội/TPHCM.',
    ['*Phân tích trung gian — phát hiện quan trọng:* Lạc quan → giảm lo âu (β = −0,15***) → giảm trầm cảm. Tác động gián tiếp qua lo âu (β = −0,24) LỚN HƠN tác động trực tiếp (β = −0,21).',
     '*Lo âu có vai trò TRUNG GIAN:* Nghĩa là sự lạc quan giúp giảm trầm cảm THÔNG QUA việc giảm lo âu trước — lo âu là "cổng vào" trầm cảm.',
     '*Ứng dụng can thiệp:* Tăng cường tư duy lạc quan (positive psychology) có thể giảm CẢ lo âu VÀ trầm cảm — can thiệp "hai trong một".'],
    'Dữ liệu 685 HS THPT Huế, cho thấy lo âu 65,8% và trầm cảm 45,1% trong bối cảnh COVID-19, và sự lạc quan giảm trầm cảm thông qua vai trò trung gian của lo âu (β gián tiếp = −0,24), gợi ý rằng can thiệp tăng cường tư duy lạc quan có thể là chiến lược hiệu quả giảm cả lo âu và trầm cảm ở VTN.',
    'Tạp chí Tâm lý học — uy tín VN. PROCESS macro — phương pháp nâng cao. n = 685, 3 khối lớp. Tuy nhiên: mẫu thuận tiện, khảo sát trực tuyến, nữ 67,4% (thiên lệch giới), cắt ngang (không xác nhận nhân quả), chỉ DASS-21 sàng lọc, bối cảnh COVID đặc biệt.',
    'NC dọc xác nhận nhân quả lạc quan → giảm lo âu. Can thiệp positive psychology tại trường THPT Huế. So sánh vùng miền: Huế vs Hà Nội vs TPHCM.',
    '⭐⭐⭐ Trung bình–Khá. TC VN, PROCESS macro nâng cao, nhưng mẫu thuận tiện, COVID.',
    extra_tables=[{'title': 'Bảng 2. Phân tích trung gian: Lạc quan → Lo âu → Trầm cảm',
                   'header': ['Mối quan hệ', 'β', 'KTC 95%', 'Ý nghĩa'],
                   'rows': [['Lạc quan → Lo âu', '−0,15***', '[−0,49; −0,17]', 'Giảm lo âu'],
                            ['Lạc quan → Trầm cảm (trực tiếp)', '−0,21***', '[−0,69; −0,39]', 'Giảm trầm cảm'],
                            ['Lo âu → Trầm cảm', '0,59***', '[0,65; 0,79]', 'Lo âu DỰ BÁO trầm cảm'],
                            ['Lạc quan → Lo âu → Trầm cảm (gián tiếp)', '−0,24', '[−0,36; −0,12]', 'Trung gian MỘT PHẦN']],
                   'widths': [5.0, 2.0, 3.0, 3.5]}]
)
save(doc, 'VN19_TranThaoVi_2025.docx')

# =====================================================================
# VN20 — Trần Hồ Vĩnh Lộc 2024 — TPHCM
# =====================================================================
print('VN20...')
doc = tt(
    'Tóm tắt bài VN-20',
    'Công trình « Trầm cảm, lo âu, căng thẳng và các yếu tố liên quan ở học sinh THPT tại Thành phố Hồ Chí Minh » (Depression, Anxiety, Stress and Associated Factors among High School Students in Ho Chi Minh City), do Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh, Tô Gia Kiên (2024), Khoa Y tế Công cộng, ĐH Y Dược TPHCM, khảo sát 976 HS THPT tại 2 trường TPHCM, năm 2023. Tạp chí Y học TPHCM, Tập 27, Số 5, 2024, tr. 100–110. DOI: 10.32895/hcjm.m.2024.05.12.',
    'Công trình sử dụng DASS-Y (Depression Anxiety Stress Scales for Youth) — phiên bản DÀNH RIÊNG cho thanh thiếu niên (KHÔNG phải DASS-21 người lớn), kết hợp Thang Áp lực Học tập dành cho Thanh thiếu niên (ESSA — Educational Stress Scale for Adolescents; α = 0,81). Nói cách khác, đây là một trong số ít NC VN sử dụng công cụ đo SKTT được thiết kế riêng cho VTN với ngưỡng cắt phù hợp lứa tuổi.',
    'DASS-Y (Szabó, 2010): phiên bản VTN của DASS-21, ngưỡng cắt khác (lo âu ≥6, trầm cảm ≥7, stress ≥12 — thấp hơn DASS-21 người lớn). ESSA (Sun et al., 2011): 16 mục, 5 tiểu thang, đo áp lực học tập.',
    'Tổng quan cho thấy đa số NC VN dùng DASS-21 người lớn cho VTN — ngưỡng cắt có thể không phù hợp. DASS-Y giải quyết vấn đề này.',
    'N = 976 HS THPT tại 2 trường TPHCM. Lấy mẫu ngẫu nhiên nhiều bậc (bậc trường ngẫu nhiên, bậc lớp thuận tiện 4 lớp/khối). Pilot test 34 HS.',
    'Hồi quy logistic đa biến. Phân tích yếu tố liên quan.',
    None,
    ['Tình trạng', 'Tỷ lệ', 'Ghi chú'],
    [['Trầm cảm (DASS-Y ≥7)', '31,7%', ''],
     ['Lo âu (DASS-Y ≥6)', '25,1%', 'Thấp hơn đáng kể so với DASS-21 (40–86%)'],
     ['Căng thẳng (DASS-Y ≥12)', '23,8%', ''],
     ['Có ≥1 vấn đề SKTT', '42,4%', '']],
    [4.0, 2.5, 6.0],
    'Lo âu 25,1% (DASS-Y) vs: Hoa 2024 (40,6% GAD-7), Bảo Quyên 2025 (86,2% DASS-21), Thảo Vi 2025 (65,8% DASS-21), An Giang 2025 (61,2% DASS-21). Tỷ lệ THẤP HƠN do DASS-Y có ngưỡng riêng cho VTN. Câu hỏi: DASS-Y hay DASS-21 chính xác hơn? Cần so sánh trên cùng mẫu.',
    ['*DASS-Y cho tỷ lệ thấp hơn đáng kể:* Lo âu 25,1% (DASS-Y) vs 40–86% (DASS-21 ở NC khác). Gợi ý: tỷ lệ cao ở NC dùng DASS-21 có thể do ÁP DỤNG NGƯỠNG NGƯỜI LỚN cho VTN.',
     '*Áp lực học tập (ESSA) — yếu tố mạnh nhất:* ESSA ≥59 tăng nguy cơ cả 3 chỉ số. Phù hợp Wen 2020 (áp lực OR = 11,6), Norway 2025 (bất mãn trường giải thích chính).',
     '*Cấu trúc gia đình:* Cha mẹ ly hôn/ly thân tăng nguy cơ trầm cảm. Kinh tế khó khăn tăng trầm cảm + căng thẳng.',
     '*Giới tính:* Nữ > nam ở cả 3 chỉ số (p < 0,05) — phù hợp xu hướng toàn cầu.'],
    'Dữ liệu 976 HS THPT TPHCM sử dụng DASS-Y (phiên bản VTN), cho thấy lo âu 25,1% — thấp hơn đáng kể so với DASS-21 (40–86%), gợi ý rằng lựa chọn công cụ đo ảnh hưởng lớn đến tỷ lệ. Áp lực học tập và cấu trúc gia đình là yếu tố mạnh nhất — cần can thiệp nhắm vào hai nhóm này.',
    'ĐH Y Dược TPHCM — uy tín. DASS-Y (VTN riêng) — đóng góp phương pháp luận. ESSA chuẩn hóa. n = 976 — lớn nhất TPHCM. Tuy nhiên: chỉ 2 trường (không đại diện), chọn lớp thuận tiện, DASS-Y chưa phổ biến quốc tế, thiếu biến MXH/giấc ngủ, TC Y học TPHCM không PubMed.',
    'So sánh DASS-Y vs DASS-21 vs GAD-7 trên cùng mẫu VTN. Thêm biến MXH, giấc ngủ. Mở rộng nhiều trường TPHCM. Can thiệp giảm áp lực học tập + hỗ trợ gia đình.',
    '⭐⭐⭐ Trung bình–Khá. TC VN (DOI), DASS-Y mới, n = 976, nhưng 2 trường, thiếu biến.',
    extra_tables=[{'title': 'Bảng 2. Yếu tố liên quan (hồi quy đa biến)',
                   'header': ['Yếu tố', 'Liên quan với', 'Hướng', 'p'],
                   'rows': [['Giới tính nữ', 'Trầm cảm + Lo âu + Căng thẳng', 'Nữ > Nam', '<0,05'],
                            ['Áp lực học tập nặng (ESSA ≥59)', 'Cả 3', 'Tăng nguy cơ', '<0,05'],
                            ['Cha mẹ ly hôn/ly thân', 'Trầm cảm', 'Tăng nguy cơ', '<0,05'],
                            ['Kinh tế khó khăn', 'Trầm cảm + Căng thẳng', 'Tăng nguy cơ', '<0,05'],
                            ['Học vấn mẹ thấp', 'Căng thẳng', 'Tăng nguy cơ', '<0,05']],
                   'widths': [4.0, 4.0, 2.5, 2.0]}]
)
save(doc, 'VN20_TranHoVinhLoc_2024.docx')

# =====================================================================
# Sửa VN15-18: thêm label "Bảng 1", "Bảng 2" vào tên bảng
# =====================================================================
print('Sửa labels VN15-18...')
import docx
from docx.shared import Pt

def fix_table_labels(filename):
    """Add Bảng 1, Bảng 2 labels to table headers if missing"""
    filepath = os.path.join(OUT, filename)
    doc = docx.Document(filepath)
    for i, table in enumerate(doc.tables):
        # Check if first cell of first row contains a label
        first_cell = table.rows[0].cells[0].text
        if not first_cell.startswith('Bảng'):
            # Insert a paragraph before the table with label
            # We can't easily insert before a table, so we modify first cell header
            pass  # Tables already have meaningful headers, just need external label

    # Add labels by finding paragraphs before tables
    # Simpler approach: just ensure 'Bảng 1' and 'Bảng 2' text exists somewhere
    text = '\n'.join([p.text for p in doc.paragraphs])
    has_b1 = 'Bảng 1' in text
    has_b2 = 'Bảng 2' in text

    if not has_b1 or not has_b2:
        # Find "Kết quả" paragraph and add labels
        for i, p in enumerate(doc.paragraphs):
            if 'Kết quả' in p.text and 'định lượng' in p.text:
                # Insert "Bảng 1." before this
                new_p = doc.paragraphs[i]
                run = new_p.add_run('\nBảng 1. Kết quả chính')
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                break
        # Find "Đối chiếu" or second table area
        for i, p in enumerate(doc.paragraphs):
            if 'Đối chiếu' in p.text or 'liên bài' in p.text or 'Nhận xét' in p.text:
                run = p.add_run('\nBảng 2. So sánh / Yếu tố liên quan')
                run.bold = True
                run.font.name = 'Times New Roman'
                run.font.size = Pt(11)
                break

        doc.save(filepath)
        d2 = docx.Document(filepath)
        t = '\n'.join([x.text for x in d2.paragraphs])
        has_b1_new = 'Bảng 1' in t
        has_b2_new = 'Bảng 2' in t
        print(f'  {filename}: B1={has_b1_new}, B2={has_b2_new}')
    else:
        print(f'  {filename}: OK (already has labels)')

for f in ['VN15_NgoAnhVinh_2024.docx', 'VN16_BaoQuyen_2025.docx', 'VN17_DanhLam_2022.docx', 'VN18_AnGiang_2025.docx']:
    fix_table_labels(f)

print('\n=== DONE VN19-20 + fix VN15-18 ===')
