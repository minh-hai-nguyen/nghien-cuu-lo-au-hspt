# -*- coding: utf-8 -*-
"""Tu dien form Buoc 5 - Giay tiep nhan PBDL lan 1 cho NCS Cong Thi Hang.
- Doc template goc + giu nguyen format
- Dien thong tin Lan 1 (Table 1 + Para 4-15 + Table 3)
- Giu Lan 2 Case 1 + Case 2 nguyen ven (template phong sau)
29/05/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from copy import deepcopy

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LATS = os.path.join(ROOT, 'Luận án TS')

# Find template via os.listdir (Unicode-safe)
TEMPLATE = None
for f in os.listdir(LATS):
    if '5.' in f and 'Gi' in f and 'Ph' in f and 'b' in f and 'doc' in f.lower():
        TEMPLATE = os.path.join(LATS, f)
        break

print(f'Template: {TEMPLATE}')
OUT = os.path.join(LATS, 'Buoc5_GiayTiepNhan_PBDL_Lan1_DaDien_29052026.docx')

d = Document(TEMPLATE)


def replace_in_paragraph(para, old_substr, new_text, preserve_format=True):
    """Thay the substring trong paragraph, giu format tu run dau tien."""
    full = para.text
    if old_substr not in full:
        return False
    new_full = full.replace(old_substr, new_text)
    # Get format from first run
    if para.runs:
        ref_run = para.runs[0]
        font_name = ref_run.font.name
        font_size = ref_run.font.size
        bold = ref_run.bold
        italic = ref_run.italic
    else:
        font_name, font_size, bold, italic = 'Times New Roman', Pt(13), None, None

    # Clear paragraph
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    # Add single new run with replaced text
    new_run = para.add_run(new_full)
    new_run.font.name = font_name
    if font_size:
        new_run.font.size = font_size
    new_run.bold = bold
    new_run.italic = italic
    return True


def replace_cell_text(cell, old_substr, new_text):
    """Thay text trong cell."""
    for para in cell.paragraphs:
        if old_substr in para.text:
            replace_in_paragraph(para, old_substr, new_text)
            return True
    return False


def set_para_text(para, new_text, font_size_pt=13, bold=False, italic=False):
    """Replace toan bo paragraph bang text moi."""
    if para.runs:
        ref_run = para.runs[0]
        font_name = ref_run.font.name or 'Times New Roman'
    else:
        font_name = 'Times New Roman'
    for r in list(para.runs):
        r._element.getparent().remove(r._element)
    nr = para.add_run(new_text)
    nr.font.name = font_name
    nr.font.size = Pt(font_size_pt)
    nr.bold = bold
    nr.italic = italic


# ============================================================
# DIEN LAN 1
# ============================================================

# --- Table 1: Header (date) ---
# R0C1 chua: "Hà Nội, ngày........ tháng........năm"
t1 = d.tables[0]
header_right = t1.rows[0].cells[1]
for para in header_right.paragraphs:
    if 'ngày' in para.text and 'tháng' in para.text:
        replace_in_paragraph(para, 'ngày........ tháng........năm',
                             'ngày 29 tháng 5 năm 2026')
        break
print('[Lan 1] Table 1: header date filled')


# --- Paragraphs Lan 1 (index 0-based: para [1]-[15]) ---
paragraphs = d.paragraphs

# Para 4 (1-based): "Bộ phận tiếp nhận và trả kết quả tiếp nhận hồ sơ của: ………"
# → "...của: Nghiên cứu sinh Công Thị Hằng"
for para in paragraphs[:20]:
    if 'Bộ phận tiếp nhận' in para.text and 'hồ sơ của' in para.text:
        replace_in_paragraph(para, '…………...………………………',
                             'Nghiên cứu sinh Công Thị Hằng')
        print(f'[Lan 1] Filled: NCS name')
        break

# Para 5: "Địa chỉ: ………"
for para in paragraphs[:20]:
    if para.text.startswith('Địa chỉ:') and '………' in para.text:
        replace_in_paragraph(para, '……………………………………………………………………………….',
                             '[NCS điền địa chỉ liên lạc]')
        print(f'[Lan 1] Filled: address placeholder')
        break

# Para 6: "Số điện thoại: ………"
for para in paragraphs[:20]:
    if para.text.startswith('Số điện thoại:') and '………' in para.text:
        # add email field too
        set_para_text(para,
                      'Số điện thoại: [NCS điền SĐT]   Email: [NCS điền email]',
                      font_size_pt=13)
        print(f'[Lan 1] Filled: phone + email placeholders')
        break

# Para 10: "2. Số lượng hồ sơ: …….(bộ)" → "02 (bộ)"
for para in paragraphs[:20]:
    if para.text.startswith('2. Số lượng hồ sơ:'):
        replace_in_paragraph(para, '…….(bộ)', '02 (bộ)')
        print(f'[Lan 1] Filled: so luong ho so = 02')
        break

# Para 11: "3. Thời gian nhận hồ sơ: …" → leave (school fills upon receipt)
# Para 15: "Vào Sổ theo dõi hồ sơ" → leave (school fills)


# --- Table 3: Signature row ---
# R0C0: "Người nộp hồ sơ (Ký và ghi rõ họ tên)" → add NCS name
t3 = d.tables[2]
sig_left = t3.rows[0].cells[0]
# Add NCS name as a new paragraph at the bottom
sig_para = sig_left.add_paragraph()
sig_para.alignment = 1  # CENTER
nr = sig_para.add_run('\n\n\n')
nr.font.name = 'Times New Roman'
nr.font.size = Pt(13)
sig_para2 = sig_left.add_paragraph()
sig_para2.alignment = 1
nr2 = sig_para2.add_run('Công Thị Hằng')
nr2.font.name = 'Times New Roman'
nr2.font.size = Pt(13)
nr2.bold = True
print('[Lan 1] Table 3: NCS name added to signature box')


# --- Table 2: Docs list (8 items pre-filled in template — verify gi chu) ---
# Em add cot Ghi chu "✓" cho moi item de NCS chuan bi
t2 = d.tables[1]
print(f'[Lan 1] Table 2 docs list: {len(t2.rows)-1} items (pre-filled in template)')


# ============================================================
# CLEAN METADATA
# ============================================================
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
cp.identifier = ''
cp.version = ''

d.save(OUT)
print()
print(f'Saved: {OUT}')
print(f'File size: {os.path.getsize(OUT)} bytes')
