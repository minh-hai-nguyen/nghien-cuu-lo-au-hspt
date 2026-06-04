"""Build doc: KTC the nao la rong/hep/vua/du - Chen 2023 vi du."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

BLUE = RGBColor(0, 112, 192)   # Câu trả lời + phần sửa typo
GREEN = RGBColor(54, 95, 44)
DARK = RGBColor(31, 73, 125)
RED = RGBColor(192, 0, 0)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = color
    return h

def add_para(text, bold=False, color=None, size=11, italic=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color:
        r.font.color.rgb = color
    return p

def add_blue_para(text, bold=False, size=11):
    """Toàn bộ tô xanh dành cho câu trả lời."""
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    r.font.color.rgb = BLUE
    return p

# TITLE
title = d.add_heading('Khoảng tin cậy (KTC) — thế nào là hẹp, vừa, rộng?', level=0)
for run in title.runs:
    run.font.color.rgb = DARK

sub = d.add_paragraph()
sub_r = sub.add_run('Áp dụng vào bài 7 (Chen 2023) với KTC trầm cảm 19,6–27,0% và lo âu 11,2–17,0%')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# =========================================================
# CÂU HỎI GỐC — có sửa typo, phần sửa tô xanh
# =========================================================
qbox = d.add_table(rows=1, cols=1)
qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
p = cell.paragraphs[0]
r = p.add_run('Câu hỏi của thầy (có sửa chính tả — phần sửa tô xanh):')
r.bold = True

p2 = cell.add_paragraph()
# "Trong bài 07" - chữ T thêm tô xanh
r_blue = p2.add_run('T')
r_blue.font.color.rgb = BLUE
r_blue.font.size = Pt(11)
p2.add_run('rong bài 07 có câu: "khoảng tin cậy khá rộng cho trầm cảm (19,6–27,0 %) và lo âu (11,2–17,0 %)". Vậy KTC thế nào là vừa, là đủ, là trung bình? ')
# "Em" sửa từ "Bác" tô xanh
r_blue2 = p2.add_run('Em')
r_blue2.font.color.rgb = BLUE
r_blue2.font.size = Pt(11)
p2.add_run(' viết câu trả lời vào file doc nhé.')
d.add_paragraph()

# =========================================================
# BỐI CẢNH (màu đen — background)
# =========================================================
add_heading('Bối cảnh', level=1)

add_para('Bài 7 là Chen và cộng sự (2023) đăng trên BMC Psychiatry, khảo sát 63.205 HS THCS và THPT tại thành phố Tự Cống (Trung Quốc), lấy mẫu cụm từ 2 quận và 1 huyện. Tác giả báo cáo tỷ lệ có trọng số kèm KTC 95 %:')
for it in [
    'Trầm cảm: 23,3 % [KTC 95 %: 19,6 %; 27,0 %] → rộng 7,4 điểm phần trăm',
    'Lo âu: 14,1 % [KTC 95 %: 11,2 %; 17,0 %] → rộng 5,8 điểm phần trăm',
]:
    p = d.add_para_ = d.add_paragraph(it, style='List Bullet')
    for run in p.runs:
        run.font.size = Pt(11)

add_para('Câu hỏi đặt ra: với cỡ mẫu 63.205 cực lớn, vì sao KTC vẫn rộng đến 7,4 điểm? Và ngưỡng nào được coi là hẹp/vừa/rộng?')
d.add_paragraph()

# =========================================================
# KIẾN THỨC NỀN — KTC là gì, yếu tố ảnh hưởng width
# =========================================================
add_heading('Kiến thức nền', level=1)

add_para('1. KTC width là gì?', bold=True, size=12)
add_para('Width = cận trên − cận dưới. Ví dụ KTC 19,6–27,0 % → width = 7,4 điểm phần trăm. Width nhỏ = ước tính chính xác; width lớn = ước tính mơ hồ.')
d.add_paragraph()

add_para('2. Các yếu tố ảnh hưởng width', bold=True, size=12)
fac_tbl = d.add_table(rows=5, cols=2)
fac_tbl.style = 'Table Grid'
facs = [
    ('Cỡ mẫu (n)',
     'n càng LỚN → width càng NHỎ. Công thức SE ∝ 1/√n → tăng n lên 4 lần chỉ giảm width được 2 lần.'),
    ('Độ biến thiên dữ liệu (SD)',
     'SD lớn → width lớn. Với tỷ lệ (proportion), SD tối đa khi p = 50 % → KTC rộng nhất khi tỷ lệ gần 50 %.'),
    ('Thiết kế mẫu (design effect, DEFF)',
     'Cluster sampling có DEFF > 1. Width phải nhân √DEFF. Chen 2023 DEFF có thể 5-10 do sample cluster trường.'),
    ('Mức tin cậy',
     '95 % → ±1,96·SE | 99 % → ±2,58·SE (rộng hơn ~30 %). Phần lớn dịch tễ dùng 95 %.'),
    ('Weighting (trọng số)',
     'Weight không đều → variance tăng → width tăng. Hệ số "effective sample size" giảm.'),
]
for i, (k, v) in enumerate(facs):
    row = fac_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v
d.add_paragraph()

add_para('3. Ngưỡng width phổ biến trong dịch tễ học (theo prevalence)', bold=True, size=12)
th_tbl = d.add_table(rows=5, cols=3)
th_tbl.style = 'Table Grid'
hdr = th_tbl.rows[0]
for i, h in enumerate(['Mức KTC', 'Width (điểm %)', 'Áp dụng']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

ths = [
    ('HẸP (rất chính xác)', '< 2', 'Khảo sát quốc gia lớn, simple random sampling, đủ power cho phân nhóm.'),
    ('VỪA (chấp nhận được)', '2 – 5', 'Khảo sát vùng, tỉnh, n ≥ 5.000 có weighting. Phần lớn nghiên cứu dịch tễ tốt nằm ở đây.'),
    ('RỘNG (kém chính xác)', '5 – 10', 'Khảo sát cấp huyện/trường, mẫu cụm chưa tối ưu, hoặc prevalence gần 50 %.'),
    ('RẤT RỘNG (mơ hồ)', '> 10', 'Mẫu nhỏ (<500), subgroup analysis, cluster design không được điều chỉnh. Phải diễn giải rất thận trọng.'),
]
for i, (k, w, a) in enumerate(ths):
    row = th_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FFE699')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(191, 97, 14)
    row.cells[1].text = w
    row.cells[2].text = a
d.add_paragraph()

add_para('Lưu ý: các ngưỡng trên là KINH NGHIỆM THỰC HÀNH, không phải quy tắc cứng. Có 3 cách đánh giá khác bổ sung:')
for it in [
    '① Width tương đối = width / point estimate. Ví dụ KTC 23±3% → width 6 / 23 ≈ 26 %. Tương đối < 25 % là tốt.',
    '② So sánh với "minimum important difference" (MID) — nếu width nhỏ hơn MID lâm sàng → ước tính đủ chính xác.',
    '③ Precision = 1 / SE. Precision cao = KTC hẹp.',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

add_para('4. Tại sao KTC Chen 2023 rộng dù N = 63.205?', bold=True, size=12)
add_para('Thoạt nhìn, với n = 63.205 (cực lớn), width đáng lẽ phải < 1 điểm %. Nhưng Chen báo cáo width 5,8–7,4 điểm — RẤT RỘNG. Ba lý do cộng dồn:')

reason_tbl = d.add_table(rows=3, cols=2)
reason_tbl.style = 'Table Grid'
reasons = [
    ('① Design effect (DEFF) từ lấy mẫu cụm',
     'HS cùng trường tương quan cao → DEFF có thể 5–15. Effective sample size = 63.205 / DEFF ≈ 4.000–12.000, không phải 63.205.'),
    ('② Weighting không đều',
     'Các trường khác size, tỷ lệ được chọn khác nhau. Weight chênh lớn làm variance tăng.'),
    ('③ Phân nhóm chi tiết',
     'Nếu KTC này cho 1 subgroup nhỏ (ví dụ HS THCS lớp 7 nam) thì effective n còn nhỏ hơn nữa.'),
]
for i, (k, v) in enumerate(reasons):
    row = reason_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FCE4D6')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RED
    row.cells[1].text = v

d.add_paragraph()

add_para('Kết luận kỹ thuật: KTC rộng trong Chen 2023 KHÔNG phải do mẫu thiếu — mà do thiết kế cụm. Đây là "giá phải trả" để có mẫu lớn nhanh và có thể khảo sát toàn trường. Không điều chỉnh DEFF thì sẽ báo cáo KTC HẸP GIẢ, sai lệch về phía "quá tự tin".')
d.add_paragraph()

# =========================================================
# CÂU TRẢ LỜI (TÔ XANH, GOM LẠI TRƯỚC PHỤ LỤC)
# =========================================================
add_heading('Câu trả lời', level=1, color=BLUE)

add_blue_para('Không có một con số tuyệt đối cho "KTC đủ hẹp" — nhưng trong thực hành dịch tễ, bốn mức sau là tham chiếu phổ biến:', bold=True, size=12)
d.add_paragraph()

# Bảng câu trả lời với toàn bộ nội dung tô xanh
ans_tbl = d.add_table(rows=5, cols=3)
ans_tbl.style = 'Table Grid'
hdr = ans_tbl.rows[0]
for i, h in enumerate(['Mức', 'Width', 'Diễn giải']):
    cell = hdr.cells[i]
    shade_cell(cell, 'BDD7EE')  # nền xanh nhạt
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = BLUE

ans_data = [
    ('HẸP — rất chính xác', '< 2 điểm %', 'Ước tính đủ tin cậy để ra quyết định chính sách. Cần n lớn + design chặt.'),
    ('VỪA — chấp nhận được', '2 – 5 điểm %', 'Đủ tốt cho phần lớn nghiên cứu dịch tễ. Khi đọc paper Q1 nên thấy width ở mức này.'),
    ('RỘNG — kém chính xác', '5 – 10 điểm %', 'Như Chen 2023 (5,8 và 7,4). Kết quả có giá trị tham khảo nhưng không nên dùng trực tiếp ra chính sách.'),
    ('RẤT RỘNG — mơ hồ', '> 10 điểm %', 'Coi như "không biết rõ". Thường gặp ở subgroup phân tích hoặc mẫu quá nhỏ.'),
]
for i, (k, w, e) in enumerate(ans_data):
    row = ans_tbl.rows[i+1]
    cell0 = row.cells[0]
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = BLUE
    cell1 = row.cells[1]
    p1 = cell1.paragraphs[0]
    p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r1 = p1.add_run(w)
    r1.font.color.rgb = BLUE
    r1.bold = True
    cell2 = row.cells[2]
    p2 = cell2.paragraphs[0]
    r2 = p2.add_run(e)
    r2.font.color.rgb = BLUE

d.add_paragraph()

add_blue_para('Áp dụng cho Chen 2023:', bold=True, size=12)
add_blue_para('Trầm cảm KTC 19,6–27,0 % → width 7,4 điểm → thuộc mức RỘNG. Lo âu KTC 11,2–17,0 % → width 5,8 điểm → cũng thuộc mức RỘNG (sát ranh giới với "vừa").')
d.add_paragraph()

add_blue_para('Vì sao rộng dù n = 63.205 cực lớn:', bold=True, size=12)
add_blue_para('Do thiết kế lấy mẫu CỤM (chọn trường → toàn bộ HS) khiến "effective sample size" chỉ khoảng 4.000–12.000, không phải 63.205. HS cùng trường tương quan cao (intraclass correlation lớn), nên dù khảo sát 63k HS, độ chính xác chỉ tương đương 4–12k HS chọn ngẫu nhiên độc lập. Thêm weighting không đều làm variance tiếp tục tăng. Điều này giải thích vì sao width 5,8–7,4 — không phải vì mẫu thiếu mà vì thiết kế.')
d.add_paragraph()

add_blue_para('Khi đọc bài có KTC rộng 5–10 điểm %:', bold=True, size=12)
for it in [
    'Đọc kết quả như "ước tính CÓ THỂ nằm đâu đó trong khoảng này" — không nên quá tin vào con số điểm ở giữa.',
    'So sánh các nhóm (ví dụ nam vs nữ) chỉ nên làm khi KTC KHÔNG CHỒNG LẤN — nếu chồng → không khẳng định được khác biệt.',
    'Luôn hỏi: "Tại sao KTC rộng thế?" — cluster, subgroup nhỏ, hay prevalence gần 50 %?',
    'Nếu dùng kết quả này để lập chính sách, phải lấy CẬN TRÊN và CẬN DƯỚI làm kịch bản — không lấy điểm giữa.',
]:
    p = d.add_paragraph()
    r = p.add_run('• ' + it)
    r.font.color.rgb = BLUE
    r.font.size = Pt(11)

d.add_paragraph()

add_blue_para('Ghi nhớ 1 câu:', bold=True, size=12)
add_blue_para('"KTC < 2 điểm % là hẹp, 2–5 là vừa (chuẩn phần lớn NC Q1), 5–10 là rộng, > 10 là rất rộng. Chen 2023 rộng 5,8–7,4 điểm chủ yếu do DEFF từ cluster sampling, không phải do thiếu mẫu."', bold=True)

d.add_paragraph()

# =========================================================
# PHỤ LỤC — TÀI LIỆU THAM KHẢO
# =========================================================
add_heading('Phụ lục — Tài liệu tham khảo', level=1)
refs = [
    'Chen và cộng sự (2023). Depression and anxiety symptoms among middle and high school students in Zigong. BMC Psychiatry. [QT007 — Bài 7, nguồn câu hỏi]',
    'Altman DG, Machin D, Bryant TN, Gardner MJ. (2000). Statistics with Confidence (2nd ed.). BMJ Books. [Giáo khoa chuẩn về KTC]',
    'Lumley T. (2010). Complex Surveys: A Guide to Analysis Using R. Wiley. [Xử lý KTC trong survey cluster]',
    'Kish L. (1965). Survey Sampling. Wiley. [Khái niệm design effect và effective sample size]',
    'Rothman KJ, Greenland S, Lash TL. (2008). Modern Epidemiology (3rd ed.). Lippincott. [Chương về interval estimation]',
    'Fleiss JL, Levin B, Paik MC. (2003). Statistical Methods for Rates and Proportions (3rd ed.). Wiley. [KTC cho tỷ lệ]',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

out = '01_Bao-cao/KTC_rong_hep_vua_du_Chen_2023.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
