# -*- coding: utf-8 -*-
"""
Sửa các lỗi văn bản phát hiện qua review chi tiết từng câu (3 agent 16/05/2026).
- Bài 1: bỏ cite ma GBD [7], bỏ 2 claim không nguồn [30], bỏ claim không nguồn [22],
  GAD-7 định nghĩa 2 lần.
- Bài 2: Liu nhất quán [9], định nghĩa SMD/SUCRA, bớt over-claim, giới thiệu ClearlyMe
  ở §3.3, dịch "hybrid"/"adapt".
- Bài 03: sửa spacing "nhận thức- cảm xúc", lỗi "Sức khỏe tâm thần học", viết hoa
  giữa câu, bớt over-claim nhân quả, sửa mâu thuẫn độ tin cậy [44].
Lỗi cấu trúc BẢNG của Bài 03 (Bảng 5 nhãn sai, Bảng 8 lỗi cột, thiếu ANOVA) — KHÔNG
tự sửa, nêu cho tác giả.
"""
import os
from docx import Document

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn"
B1 = os.path.join(ROOT, "Bai1_YTNC_HSTHCS_v6_16052026.docx")
B2 = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v6_16052026.docx")
B03 = os.path.join(ROOT, "bài-03", "Công Thị Hằng_v4_đã sửa.docx")


def setp(p, t):
    if p.runs:
        p.runs[0].text = t
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(t)


def whole(doc, sub, new):
    for p in doc.paragraphs:
        if sub in p.text:
            setp(p, new)
            return True
    return False


def part(doc, sub, new):
    n = 0
    for p in doc.paragraphs:
        if sub in p.text:
            setp(p, p.text.replace(sub, new))
            n += 1
    return n


# ============ BÀI 1 ============
d = Document(B1)
res = []
# [7] bỏ cite ma GBD + DALYs
res.append(("B1 §1[7] GBD", whole(d,
    "Các báo cáo Gánh nặng bệnh tật toàn cầu (GBD) công bố gần đây",
    "Rối loạn lo âu là nhóm rối loạn sức khỏe tâm thần phổ biến nhất ở lứa tuổi vị "
    "thành niên trên phạm vi toàn cầu, với xu hướng gia tăng được ghi nhận trong thập "
    "niên qua. Gánh nặng do rối loạn lo âu ở nhóm tuổi này có xu hướng tập trung nhiều "
    "hơn tại các quốc gia thu nhập trung bình và thấp, nơi nguồn lực chăm sóc sức khỏe "
    "tâm thần học đường còn hạn chế.")))
# [22] bỏ claim không nguồn + bỏ câu kết trùng [21]
res.append(("B1 §3.1[22]", whole(d,
    "Sự dịch chuyển môn học từ tiểu học lên THCS",
    "Sự dịch chuyển môn học từ tiểu học lên trung học cơ sở — đặc biệt là sự xuất hiện "
    "đột ngột của nhiều môn chuyên sâu, hệ thống đánh giá đa môn và việc chuyển đổi giáo "
    "viên — về mặt lý thuyết cũng là một nguồn áp lực đáng kể đối với học sinh; các giai "
    "đoạn chuyển cấp thường được xem là thời điểm cao của triệu chứng lo âu học đường. "
    "Mặc dù dữ liệu Việt Nam phần lớn mang tính cắt ngang, hướng của hiệu ứng nhất quán "
    "với bằng chứng quốc tế.")))
# [30] bỏ 2 claim không nguồn -> đoạn khái niệm
res.append(("B1 §3.3[30]", whole(d,
    "Bằng chứng dọc của các nghiên cứu xã hội tại Bắc Âu",
    "Bắt nạt thể chất hiếm khi tác động đơn lẻ mà thường tương tác với các yếu tố bối "
    "cảnh khác. Khi đi kèm với mức hỗ trợ xã hội thấp, trải nghiệm bị bắt nạt có thể "
    "khuếch đại nguy cơ lo âu ở học sinh; ngược lại, một mạng lưới hỗ trợ ổn định từ "
    "bạn bè và gia đình có thể làm giảm nhẹ tác động tiêu cực này. Vì vậy, các nghiên "
    "cứu về bắt nạt cần đồng thời đo lường mức độ hỗ trợ xã hội để tránh đánh giá thiên "
    "lệch hiệu ứng.")))
# [21] GAD-7 định nghĩa 2 lần + bớt "nguyên nhân"
res.append(("B1 §3.1[21] GAD-7", part(d,
    "bằng thang Rối loạn Lo âu Tổng quát 7 mục (GAD-7) ghi nhận 40,6%",
    "bằng thang GAD-7 ghi nhận 40,6%")))
res.append(("B1 §3.1[21] nguyên nhân", part(d,
    "trong đó học tập được xác định là nguyên nhân hàng đầu",
    "trong đó học tập được xác định là yếu tố liên quan hàng đầu")))
d.save(B1)
print("BÀI 1:")
for k, v in res:
    print(f"  {k}: {v}")

# ============ BÀI 2 ============
d = Document(B2)
res = []
# [9] Liu nhất quán (preview ngắn) + Cai "khiêm tốn"->"nhỏ"
res.append(("B2 §1[9] Liu", whole(d,
    "Phân tích tổng hợp lớn của Compas và cộng sự (2017) trên 80.850",
    "Phân tích tổng hợp lớn của Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành "
    "niên xác nhận chiến lược đối phó tập trung vào vấn đề – một thành tố cốt lõi của "
    "CBT – có liên hệ ngược chiều và đáng kể với triệu chứng lo âu. Tổng quan hệ thống "
    "và phân tích mạng lưới của Liu và cộng sự (2025) trên 52 thử nghiệm lâm sàng ngẫu "
    "nhiên ở bệnh nhân rối loạn lo âu tổng quát — chủ yếu là người trưởng thành — cho "
    "thấy CBT cá nhân được xếp hạng hiệu quả cao nhất trong số các hình thức cung cấp; "
    "thứ bậc chi tiết giữa các hình thức được trình bày ở mục 3.1. Phân tích mạng lưới "
    "riêng cho rối loạn lo âu xã hội ở trẻ em – vị thành niên của Xian và cộng sự (2024) "
    "báo cáo kết quả cùng hướng, trong đó các hình thức CBT — đặc biệt là CBT qua "
    "internet — được xếp hạng hiệu quả cao nhất. Trên phương diện học đường, phân tích "
    "tổng hợp của Cai và cộng sự (2025) ghi nhận các chương trình tại trường có hiệu ứng "
    "tổng hợp nhỏ trên năng lực phục hồi tâm lý của học sinh — một cấu trúc bảo vệ có "
    "liên quan đến việc giảm nguy cơ lo âu.")))
# [23] SUCRA gloss + "khẳng định"->"củng cố" + "đạt hiệu quả cao nhất"->"nâng cao"
res.append(("B2 §3.1[23] SUCRA", part(d, "chỉ số SUCRA 71,2%",
    "chỉ số diện tích dưới đường xếp hạng tích lũy (SUCRA) 71,2%")))
res.append(("B2 §3.1[23] khẳng định", part(d,
    "khẳng định nền tảng lý thuyết của CBT",
    "củng cố cho nền tảng lý thuyết của CBT")))
res.append(("B2 §3.1[23] hiệu quả cao nhất", part(d,
    "để đạt hiệu quả cao nhất ở khu vực này",
    "nhằm nâng cao hiệu quả ở khu vực này")))
# [29] định nghĩa SMD
res.append(("B2 §3.2[29] SMD", part(d, "đạt SMD = 0,17",
    "đạt hiệu số trung bình chuẩn hóa (SMD) = 0,17")))
# [35] bỏ "hybrid" thừa
res.append(("B2 §3.3 hybrid", part(d, "nền tảng kết hợp hybrid giữa",
    "nền tảng kết hợp giữa")))
# [38] "adapt"->"điều chỉnh" + giới thiệu ClearlyMe ở §3.3
res.append(("B2 §3.3[38] adapt+ClearlyMe", part(d,
    "và cần được adapt cho lứa tuổi nhỏ hơn.",
    "và cần được điều chỉnh cho lứa tuổi nhỏ hơn. Một hướng thiết kế đáng tham khảo là "
    "ứng dụng ClearlyMe — một nền tảng CBT số được đồng thiết kế cùng vị thành niên cho "
    "trầm cảm và lo âu (Li và cộng sự, 2022).")))
d.save(B2)
print("BÀI 2:")
for k, v in res:
    print(f"  {k}: {v}")

# ============ BÀI 03 ============
d = Document(B03)
res = []
res.append(("B03 spacing nhận thức", part(d, "nhận thức- cảm xúc", "nhận thức – cảm xúc")))
res.append(("B03 spacing tâm lý", part(d, "tâm lý- giáo dục", "tâm lý – giáo dục")))
res.append(("B03 spacing nhân-quả", part(d, "nhân- quả", "nhân – quả")))
res.append(("B03 'Sức khỏe tâm thần học'", part(d,
    "Sức khỏe tâm thần học ngày càng được quan tâm",
    "Sức khỏe tâm thần ngày càng được quan tâm")))
res.append(("B03 viết hoa [25]", part(d,
    "vị thành niên, Rối loạn lo âu thường biểu hiện",
    "vị thành niên, rối loạn lo âu thường biểu hiện")))
res.append(("B03 viết hoa [44]", part(d,
    "thang đo Rối loạn lo âu đạt độ tin cậy tốt",
    "thang đo rối loạn lo âu đạt độ tin cậy tốt")))
res.append(("B03 over-claim [62]", part(d,
    "Điều này cho thấy vai trò của trải nghiệm tương tác gia đình có tác dụng tốt trong can thiệp lo âu học đường.",
    "Kết quả này gợi ý rằng trải nghiệm tương tác gia đình có thể là một hướng đáng quan tâm cho các chương trình can thiệp lo âu học đường.")))
res.append(("B03 over-claim [68]", part(d,
    "yếu tố tác động không có ý nghĩa thống kê",
    "yếu tố dự báo không có ý nghĩa thống kê")))
res.append(("B03 over-claim [74]", part(d,
    "chịu ảnh hưởng đáng kể từ bối cảnh môi trường gia đình",
    "có mối liên hệ với bối cảnh môi trường gia đình")))
res.append(("B03 mâu thuẫn độ tin cậy [44]", part(d,
    "Thang đo Hoạt động gia đình và Quan hệ gia đình đạt độ tin cậy khá tốt, đáp ứng yêu cầu về độ tin cậy trong nghiên cứu tâm lý – giáo dục.",
    "Thang đo Quan hệ gia đình đạt độ tin cậy chấp nhận được; thang đo Hoạt động gia đình có hệ số tin cậy ở mức thấp, do đó các kết quả liên quan đến thang đo này cần được diễn giải một cách thận trọng.")))
d.save(B03)
print("BÀI 03:")
for k, v in res:
    print(f"  {k}: {v}")
print("[DONE]")
