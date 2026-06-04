# -*- coding: utf-8 -*-
"""Fix LA v3_1 -> v3_2: cac loi da phat hien.
27/05/2026.

Loi can sua:
A. VN017 Lam 2022:
   - Para 267: "vùng bán đô thị" -> "huyện bán nông thôn"
   - Para 268: "7,7% nhẹ" -> "11,2% nhẹ"; "24,5% vừa" -> "25,1% vừa"
   - Para 269: "dân cư bán đô thị" -> "dân cư bán nông thôn"

B. VN018 An Giang 2025:
   - Para 271 (in-text): "Lê Minh T., Nguyễn Đăng K., Ngô Anh V."
     -> "Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh"
   - Para 315 (tỷ lệ list): "Lê Minh T., 2025" -> "Nguyễn Đăng Khoa và cs., 2025"
   - Para 1368 (TLTK): "Lê Minh Thi, Nguyễn Đăng Khoa, & Ngô Anh Vinh"
     -> "Nguyễn Đăng Khoa, Lê Minh Thi, & Ngô Anh Vinh"

C. Thuat ngu DSM-5:
   - "rối loạn lo âu tổng quát" -> "rối loạn lo âu lan tỏa"
   - "lo âu tổng quát" -> "lo âu lan tỏa"
   - "RLLATQ" -> "RLLALT"
   - "Rối loạn lo âu tổng quát" -> "Rối loạn lo âu lan tỏa" (capitalized)
"""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_26052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_2_FixSoLieu_27052026.docx')

RED = RGBColor(192, 0, 0)

# Replacement list - text replacements applied globally
# Each item: (old_text, new_text, is_high_priority_fact)
TEXT_REPLACEMENTS = [
    # A. VN017 Lam 2022 - dia ly + so lieu
    ('vùng bán đô thị, gần thành phố Thanh Hóa', 'huyện bán nông thôn, gần thành phố Thanh Hóa'),
    ('dân cư bán đô thị như Yên Định', 'dân cư bán nông thôn như Yên Định'),
    # Lâm 2022 numbers
    ('có 49% học sinh lo âu, trong đó có 7,7% loại nhẹ, 24,5% loại vừa',
     'có 49,0% học sinh lo âu, trong đó có 11,2% loại nhẹ, 25,1% loại vừa'),

    # B. VN018 An Giang 2025 - thu tu tac gia (in-text)
    ('Công trình của Lê Minh T., Nguyễn Đăng K., Ngô Anh V. (2025)',
     'Công trình của Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh (2025)'),
    # Citation in para 315
    ('61,2% theo Lê Minh T., 2025', '61,2% theo Nguyễn Đăng Khoa và cs., 2025'),
    # TLTK para 1368
    ('Lê Minh Thi, Nguyễn Đăng Khoa, & Ngô Anh Vinh (2025)',
     'Nguyễn Đăng Khoa, Lê Minh Thi, & Ngô Anh Vinh (2025)'),

    # C. Thuat ngu DSM-5: tong quat -> lan toa
    # Order: longer/more-specific first to avoid double-replace
    ('Rối loạn lo âu tổng quát', 'Rối loạn lo âu lan tỏa'),
    ('rối loạn lo âu tổng quát', 'rối loạn lo âu lan tỏa'),
    ('Lo âu tổng quát', 'Lo âu lan tỏa'),
    ('lo âu tổng quát', 'lo âu lan tỏa'),
    # Abbreviations
    ('RLLATQ', 'RLLALT'),
]


def apply_replacements_in_paragraph(p, replacements, applied_tracker, para_idx):
    """Apply MULTIPLE replacements in a single paragraph.
    Builds segments from original p.text, then rebuilds runs."""
    full = p.text
    if not full:
        return False
    # Find all matches with positions
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            idx = full.find(old, cursor)
            if idx == -1: break
            matches.append((idx, idx + len(old), old, new))
            cursor = idx + len(old)
    if not matches:
        return False
    # Sort by position, then by length (longer first to win conflicts)
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    # De-overlap
    safe_matches = []
    last_end = -1
    for m in matches:
        if m[0] >= last_end:
            safe_matches.append(m)
            last_end = m[1]
    if not safe_matches:
        return False
    # Build segments
    segments = []  # (text, is_red)
    cursor = 0
    for start, end, old, new in safe_matches:
        if cursor < start:
            segments.append((full[cursor:start], False))
        segments.append((new, True))
        applied_tracker.append(f"para {para_idx}: '{old[:60]}' -> '{new[:60]}'")
        cursor = end
    if cursor < len(full):
        segments.append((full[cursor:], False))
    # Save original alignment etc. by keeping paragraph object
    # Clear all runs
    for r in p.runs:
        r.text = ''
    # Add segments as new runs
    for text, is_red in segments:
        if not text:
            continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


def apply_to_cell(cell, replacements, applied_tracker, label):
    """Apply replacements to a table cell."""
    changed = False
    for ci, p in enumerate(cell.paragraphs):
        if apply_replacements_in_paragraph(p, replacements, applied_tracker, label + f' p{ci}'):
            changed = True
    return changed


def main():
    print(f"=== Fix LA v3_1 -> v3_2 ===")
    print(f"SRC: {SRC}")
    print(f"OUT: {OUT}")
    if not os.path.exists(SRC):
        print(f"!!! SRC khong ton tai: {SRC}")
        return
    # Backup
    backup = SRC.replace('.docx', '_BACKUP_27052026.docx')
    if not os.path.exists(backup):
        shutil.copy(SRC, backup)
        print(f"Backup: {backup}")

    doc = Document(SRC)
    applied = []

    # Body paragraphs
    print(f"\n--- Fix body paragraphs ({len(doc.paragraphs)} paragraphs) ---")
    body_count = 0
    for i, p in enumerate(doc.paragraphs):
        if apply_replacements_in_paragraph(p, TEXT_REPLACEMENTS, applied, i):
            body_count += 1
    print(f"Fixed {body_count} paragraphs.")

    # Tables
    print(f"\n--- Fix tables ({len(doc.tables)} tables) ---")
    tbl_count = 0
    for ti, tb in enumerate(doc.tables):
        for ri, row in enumerate(tb.rows):
            for ci, cell in enumerate(row.cells):
                if apply_to_cell(cell, TEXT_REPLACEMENTS, applied,
                                 f'T{ti+1}r{ri}c{ci}'):
                    tbl_count += 1
    print(f"Fixed {tbl_count} cells.")

    # Save
    doc.save(OUT)
    print(f"\nSaved: {OUT}")
    print(f"Total replacements applied: {len(applied)}")
    print(f"\n--- First 60 replacements applied ---")
    for a in applied[:60]:
        print(f"  ✓ {a}")
    if len(applied) > 60:
        print(f"  ... and {len(applied)-60} more")

if __name__ == '__main__':
    main()
