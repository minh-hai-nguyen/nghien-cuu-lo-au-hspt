"""Fix 3 wrong numbers in 6 docs (auto find-replace in paragraphs + tables).
- 28,642 → 45,984  (F giới × RLLAXH)
- 45,484 → 44,484  (F giới × RLLATQ)
- 19,86  → 19,46   (M lớp 9 RLLAC)

Save with suffix _FIXED to keep original for review.
"""
import sys, io, os, shutil
from pathlib import Path
from docx import Document

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

REPLACEMENTS = [
    ('28,642', '45,984'),
    ('45,484', '44,484'),
    ('19,86', '19,46'),
]

DOCS = [
    'BO_SUNG_AI_chuong_3_luan_an_07052026.docx',
    'Binh_luan_4_bang_so_lieu_anh_07052026.docx',
    'CLEAN_chuong_3_bo_sung_de_cuong.docx',
    'CLEAN_v2_CTH_v6_chuong_4_bo_sung.docx',
    'CLEAN_v3_CTH_v6_chuong_4_bo_sung_full.docx',
    'VERIFY_REPORT_phien_07052026.docx',
]

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao')

def replace_in_runs(paragraph, replacements):
    """Replace text in runs, preserving formatting."""
    count = 0
    for run in paragraph.runs:
        text = run.text
        for old, new in replacements:
            if old in text:
                text = text.replace(old, new)
                count += 1
        if text != run.text:
            run.text = text
    return count

def fix_doc(src_path, dst_path, replacements):
    doc = Document(src_path)
    total_fixes = 0

    # Fix in paragraphs
    for p in doc.paragraphs:
        total_fixes += replace_in_runs(p, replacements)

    # Fix in tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    total_fixes += replace_in_runs(p, replacements)

    doc.save(dst_path)
    return total_fixes

def main():
    print('=== AUTO-FIX 3 SỐ SAI TRONG 6 DOCS ===\n')
    print('Replacements:')
    for old, new in REPLACEMENTS:
        print(f'  {old} → {new}')
    print()

    total_grand = 0
    for fname in DOCS:
        src = BASE / fname
        if not src.exists():
            print(f'⚠ NOT FOUND: {fname}')
            continue
        # Backup nguyên gốc với suffix _ORIG
        backup = BASE / (fname.replace('.docx', '_ORIG_BACKUP.docx'))
        if not backup.exists():
            shutil.copy(src, backup)
        # Sửa file gốc
        n = fix_doc(src, src, REPLACEMENTS)
        total_grand += n
        print(f'✓ {fname}: {n} thay thế')

    print()
    print(f'=== TỔNG: {total_grand} số liệu đã sửa trong {len(DOCS)} files ===')
    print('(Backup gốc có suffix _ORIG_BACKUP.docx)')

if __name__ == '__main__':
    main()
