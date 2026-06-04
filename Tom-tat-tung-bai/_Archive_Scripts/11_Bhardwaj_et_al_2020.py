# -*- coding: utf-8 -*-
"""Bài 11: Bhardwaj et al. 2020 — CTH v5 (≤2,5 trang)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
doc = Document()
style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.5
for s in doc.sections: s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)

def rb(t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

bl('Tóm tắt bài 11', bold=True)

rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình \u00ab Nghiên cứu mô tả đánh giá trầm cảm, lo âu và căng thẳng ở học sinh trung học phổ thông các trường công lập tại Chandigarh, Ấn Độ \u00bb (A Descriptive Study to Assess Depression, Anxiety & Stress among Higher Secondary Students of Government Schools of Chandigarh, India), do Ridhul Bhardwaj và cộng sự (2020), khảo sát trên 288 học sinh (13\u201321 tuổi, 72,9% nhóm 16\u201318) tại 5 trường công lập ở Chandigarh, xuất bản trên Journal of IPHA Chandigarh State Branch.')

rb('Phương pháp nghiên cứu')
bl('Công trình sử dụng DASS-21 (Thang đo Trầm cảm Lo âu Căng thẳng 21 mục, Lovibond & Lovibond, 1995) để đánh giá đồng thời ba tình trạng. Nói cách khác, nghiên cứu mô tả cắt ngang sử dụng lấy mẫu xác suất tỷ lệ (probability proportionate sampling) từ 5 trường.', bold=True)
bl('Pilot test trên 20 học sinh. Phân tích bằng chi-square, tương quan Pearson, ANOVA và t-test. Chấp thuận đạo đức đã được lấy.')

rb('Kết quả nghiên cứu định lượng')
bl('Tỷ lệ trầm cảm, lo âu, căng thẳng theo DASS-21 (N = 288):', bold=True)
tbl(
    ['Tình trạng', 'Tỷ lệ tổng', 'Điểm TB', 'Nặng+Cực nặng'],
    [
        ['Lo âu', '81,9%', '14,09 \u00b1 6,7', '46,8%'],
        ['Trầm cảm', '64,9%', '13,05 \u00b1 7,2', '16,3%'],
        ['Căng thẳng', '55,2%', '15,6 \u00b1 7,1', '8,6%'],
    ],
    widths=[3.0, 2.5, 3.0, 3.0]
)
doc.add_paragraph()

bl('Mối liên quan giữa ba tình trạng:', bold=True)
tbl(
    ['Cặp liên quan', 'OR', 'Tương quan Pearson r'],
    [
        ['Căng thẳng \u2194 Trầm cảm', '17,26', '0,664'],
        ['Lo âu \u2194 Trầm cảm', '10,93', '0,617'],
        ['Lo âu \u2194 Căng thẳng', '9,77', '0,575'],
    ],
    widths=[4.5, 2.5, 4.0]
)
doc.add_paragraph()

bh3('Khác biệt giới tính : Nữ có tỷ lệ lo âu nặng + cực nặng 63,7% so với nam 36,0% (P < 0,001) \u2014 chênh lệch lớn nhất trong 11 bài. Trầm cảm: nữ 23,9% nặng vs nam 11,4% (P = 0,04). Căng thẳng: nữ 44,2% trung bình/nặng vs nam 23,3% (P < 0,001).')

rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bl('*Tỷ lệ lo âu CAO NHẤT trong 11 bài* (81,9%). DASS-21 với ngưỡng cắt thấp (≥8 cho lo âu) dẫn đến tỷ lệ rất cao khi bao gồm cả mức "nhẹ". Nếu chỉ tính nặng + cực nặng (46,8%), tỷ lệ vẫn rất đáng báo động.')
bl('*Đồng mắc rất cao.* OR giữa các cặp rối loạn từ 9,77 đến 17,26, xác nhận sự chồng chéo mạnh giữa trầm cảm, lo âu và căng thẳng \u2014 gợi ý sàng lọc cả 3 tình trạng đồng thời.')

rb('Kết luận')
bl('Lo âu 81,9%, trầm cảm 64,9%, căng thẳng 55,2% ở 288 học sinh trường công Chandigarh \u2014 tỷ lệ cao nhất trong 11 bài. Nữ bị ảnh hưởng nặng hơn nam ở cả 3 tình trạng. Đồng mắc rất cao (OR = 9,77\u201317,26). Cần tăng cường nhận thức của giáo viên và phụ huynh về nguồn căng thẳng ngoài học tập, và mở rộng đánh giá SKTT ra nhiều lĩnh vực.', bold=True)

rh2('Phản biện')
bl('Tạp chí Journal of IPHA Chandigarh State Branch là tạp chí cấp chi nhánh địa phương, không có impact factor, không được lập chỉ mục PubMed/Scopus. Cỡ mẫu nhỏ (288), chỉ 5 trường công lập. Tỷ lệ 81,9% lo âu cao bất thường \u2014 có thể do ngưỡng cắt thấp hoặc đặc thù mẫu. Chỉ dùng chi-square đơn biến, không có hồi quy đa biến (AOR). Pilot test chỉ 20 học sinh.')

rh2('Hướng nghiên cứu tiếp theo')
bl('Tái lặp với cỡ mẫu lớn hơn, bao gồm cả trường tư thục. Sử dụng hồi quy đa biến để xác định yếu tố nguy cơ độc lập. So sánh DASS-21 với công cụ chẩn đoán trên cùng mẫu. Xuất bản trên tạp chí có chỉ mục quốc tế.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50 Thấp-Trung bình. DASS-21 xác thực nhưng tạp chí không chỉ mục, cỡ mẫu nhỏ, thiếu phân tích đa biến.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('11_Bhardwaj_et_al_2020.docx')
print('11_Bhardwaj_et_al_2020.docx OK')
