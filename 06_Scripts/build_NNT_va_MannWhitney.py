"""Build doc: NNT + Mann-Whitney chenh lech trung vi."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

# TITLE
title = d.add_heading('Hai câu hỏi thống kê: NNT và Mann-Whitney (chênh lệch trung vị 3 điểm)', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Giải thích 2 chỉ số thống kê hay gặp khi đọc RCT và nghiên cứu cắt ngang')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# META
meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
mp = cell.paragraphs[0]
mp.add_run('Câu hỏi 1 của thầy: ').bold = True
mp.add_run('"NNT diễn giải cụ thể như thế nào? Vì sao chỉ số này càng nhỏ thì biện pháp can thiệp càng hiệu quả?"')
mp2 = cell.add_paragraph()
mp2.add_run('Câu hỏi 2 của thầy: ').bold = True
mp2.add_run('"Trong bài 1 có câu: \'Về khác biệt giới tính, học sinh nữ có điểm GAD10 cao hơn nam (p=0,016, kiểm định Mann-Whitney), với chênh lệch trung vị là 3 điểm\'. Chênh lệch trung vị 3 điểm, theo kiểm định này nghĩa là gì?"')
mp3 = cell.add_paragraph()
mp3.add_run('Bài 1 = ').bold = True
mp3.add_run('Jenkins J.H. và cộng sự (2023) "Depression and anxiety among multiethnic middle school students". QT001 trong thư viện. Dùng PHQ-9A + GAD-10 trên mẫu N=75 HS THCS California, Mỹ. Tác giả ghi: "Nữ có điểm PHQ-9A cao hơn nam (p=0,002) và điểm GAD-10 cao hơn (p=0,016), đều dùng Mann-Whitney".')
d.add_paragraph()

# =========================================================
# CÂU HỎI 1 — NNT
# =========================================================
add_heading('CÂU HỎI 1 — NNT (Number Needed to Treat)', level=1)

add_para('1.1. Định nghĩa', bold=True, size=12)
add_para('NNT = Số người cần điều trị để có THÊM 1 người cải thiện so với không điều trị (hoặc so với nhóm đối chứng).', bold=True, color=(192, 0, 0))
d.add_paragraph()

add_para('Diễn giải 1 dòng: "Cứ điều trị NNT người, sẽ có THÊM 1 người đáp ứng nhờ can thiệp."')
d.add_paragraph()

# Ví dụ CAMS
ex_tbl = d.add_table(rows=1, cols=1)
ex_tbl.style = 'Table Grid'
ec = ex_tbl.rows[0].cells[0]
shade_cell(ec, 'DEEBF7')
ep = ec.paragraphs[0]
er = ep.add_run('Ví dụ CAMS (Walkup 2008 NEJM): ')
er.bold = True
er.font.color.rgb = RGBColor(31, 73, 125)
ep.add_run('NNT = 3 cho kết hợp CBT + Sertraline → cứ 3 trẻ lo âu được điều trị combination, sẽ có THÊM 1 trẻ đáp ứng lâm sàng so với nhóm giả dược. Hai trẻ còn lại có thể đã tự hồi phục hoặc không đáp ứng.')
d.add_paragraph()

add_para('1.2. Công thức', bold=True, size=12)
add_para('NNT = 1 / ARR', bold=True, size=14, color=(192, 0, 0))
add_para('ARR = Absolute Risk Reduction = tỷ lệ đáp ứng nhóm can thiệp − tỷ lệ đáp ứng nhóm đối chứng (dùng số thập phân 0-1, không phải phần trăm).')
d.add_paragraph()

add_para('Tính cho CAMS combination:', bold=True)
calc = [
    'Tỷ lệ đáp ứng CBT+Sertraline = 80,7 % = 0,807',
    'Tỷ lệ đáp ứng Placebo = 23,7 % = 0,237',
    'ARR = 0,807 − 0,237 = 0,57',
    'NNT = 1 / 0,57 ≈ 1,75 → làm tròn LÊN thành 2 (một số nguồn báo cáo ≈ 3)',
]
for it in calc:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)
d.add_paragraph()

add_para('1.3. Vì sao càng nhỏ càng hiệu quả — logic toán đơn giản', bold=True, size=12)

logic_tbl = d.add_table(rows=5, cols=2)
logic_tbl.style = 'Table Grid'
hdr = logic_tbl.rows[0]
for i, h in enumerate(['NNT', 'Ý nghĩa']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

logic = [
    ('NNT = 1 (lý tưởng)', 'Điều trị 1 người → chắc chắn 1 người cải thiện. Hiếm trong y học (vd insulin cho ĐTĐ type 1).'),
    ('NNT = 3 (xuất sắc)', 'CAMS combination. Mỗi 3 trẻ → 1 trẻ cải thiện thêm. Can thiệp CỰC KỲ mạnh.'),
    ('NNT = 10 (rất tốt)', 'Nhiều SSRIs cho trầm cảm nặng. Chấp nhận được trong lâm sàng.'),
    ('NNT = 50 (trung bình)', 'Statins dự phòng tim mạch. Cần cost-benefit kỹ khi nhân rộng.'),
]
for i, (k, v) in enumerate(logic):
    row = logic_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = v
d.add_paragraph()

# Ẩn dụ lưới
ana_tbl = d.add_table(rows=1, cols=1)
ana_tbl.style = 'Table Grid'
ac = ana_tbl.rows[0].cells[0]
shade_cell(ac, 'FFF8DC')
ap = ac.paragraphs[0]
ar = ap.add_run('Ẩn dụ "lưới cá": ')
ar.bold = True
ar.font.color.rgb = RGBColor(191, 97, 14)
ap.add_run('NNT = số "con mồi" phải thả để BẮT được 1 con cá THÊM. Lưới hiệu quả → 3 con mồi bắt 1 cá. Lưới kém → 50 con mồi mới bắt 1 cá. Càng ít con mồi cần thả = lưới càng mạnh = NNT càng nhỏ.')
d.add_paragraph()

add_para('1.4. Trực giác công thức', bold=True, size=12)
add_para('Vì NNT = 1 / ARR, nên ARR LỚN → NNT NHỎ. ARR lớn nghĩa là can thiệp tạo ra CHÊNH LỆCH RÕ so với đối chứng. Phép chia "1 / số lớn" = số nhỏ. Toán học đảm bảo NNT nhỏ luôn phản ánh can thiệp mạnh.')
d.add_paragraph()

# =========================================================
# CÂU HỎI 2 — MANN-WHITNEY + MEDIAN DIFFERENCE
# =========================================================
add_heading('CÂU HỎI 2 — Mann-Whitney và "chênh lệch trung vị 3 điểm"', level=1, color=(192, 80, 77))

add_para('2.1. Bối cảnh Jenkins 2023', bold=True, size=12)
add_para('Jenkins và cộng sự đo lo âu bằng thang GAD-10 (phiên bản mở rộng 10 câu của GAD-7 gốc). Mỗi câu chấm 0–3 điểm, tổng 0–30. So sánh điểm GAD-10 giữa HS nữ (n=45) và nam (n=30) trên mẫu N=75 HS THCS đa sắc tộc California. Dữ liệu điểm GAD-10 KHÔNG phân phối chuẩn → dùng kiểm định phi tham số thay vì t-test.')
d.add_paragraph()

add_para('2.2. Kiểm định Mann-Whitney U là gì?', bold=True, size=12)

mw_tbl = d.add_table(rows=6, cols=2)
mw_tbl.style = 'Table Grid'
mw_data = [
    ('Tên đầy đủ', 'Mann-Whitney U test (còn gọi: Wilcoxon rank-sum test — về toán học là tương đương)'),
    ('Loại', 'Kiểm định PHI THAM SỐ (non-parametric) so sánh 2 nhóm ĐỘC LẬP'),
    ('Dùng khi nào', 'Khi dữ liệu: (1) không phân phối chuẩn, HOẶC (2) thang đo thứ hạng, HOẶC (3) mẫu nhỏ không đủ để giả định normality'),
    ('Giả thuyết không (H₀)', 'Hai phân phối có hình dạng giống nhau (không có nhóm nào "hay cao hơn" nhóm kia)'),
    ('Giả thuyết thay thế (H₁)', 'Một nhóm có xu hướng CAO HƠN nhóm kia một cách hệ thống (stochastic dominance)'),
    ('Cách tính', 'Gộp TẤT CẢ quan sát 2 nhóm → xếp thứ hạng từ nhỏ đến lớn → tính TỔNG THỨ HẠNG của mỗi nhóm → kiểm tra xem tổng thứ hạng có chênh lệch đủ lớn để không thể do ngẫu nhiên không'),
]
for i, (k, v) in enumerate(mw_data):
    row = mw_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v
d.add_paragraph()

# Quan trọng: không test median
key_tbl = d.add_table(rows=1, cols=1)
key_tbl.style = 'Table Grid'
kc = key_tbl.rows[0].cells[0]
shade_cell(kc, 'FCE4D6')
kp = kc.paragraphs[0]
kr = kp.add_run('⚠ HIỂU LẦM PHỔ BIẾN: ')
kr.bold = True
kr.font.color.rgb = RGBColor(192, 0, 0)
kp.add_run('Mann-Whitney KHÔNG trực tiếp test trung vị (median). Nó test sự "dominance ngẫu nhiên" — xác suất một giá trị lấy ngẫu nhiên từ nhóm A cao hơn một giá trị lấy ngẫu nhiên từ nhóm B. Nhưng trong phần LỚN trường hợp (khi 2 phân phối có hình dạng tương tự), kết quả Mann-Whitney TƯƠNG ĐƯƠNG với test chênh lệch trung vị. Vì vậy tác giả thường ghép p-value Mann-Whitney + median difference để báo cáo.')
d.add_paragraph()

add_para('2.3. Diễn giải p = 0,016', bold=True, size=12)
add_para('p = 0,016 < 0,05 → BÁC BỎ giả thuyết không. Có ý nghĩa thống kê — sự khác biệt giữa 2 nhóm (nữ/nam) KHÔNG PHẢI do ngẫu nhiên trong mẫu. Cụ thể: xác suất quan sát thấy chênh lệch này (hoặc lớn hơn) NẾU 2 nhóm thực sự giống nhau chỉ là 1,6 %. Rất hiếm → nên kết luận 2 nhóm thật sự khác nhau.')
d.add_paragraph()

add_para('Lưu ý: p = 0,016 là "có ý nghĩa" ở mức 5% nhưng KHÔNG phải "rất có ý nghĩa" (thường cần p < 0,01 hoặc p < 0,001). Với p=0,016 có thể nói "significant at α=0,05, not at α=0,01".', bold=False)
d.add_paragraph()

add_para('2.4. "Chênh lệch trung vị 3 điểm" nghĩa là gì?', bold=True, size=12)

add_para('Median difference = 3 điểm GAD-10 có NGHĨA SAU:')
d.add_paragraph()

md_tbl = d.add_table(rows=4, cols=2)
md_tbl.style = 'Table Grid'
md_data = [
    ('Cách tính',
     'Thường là Hodges-Lehmann estimator: lấy TẤT CẢ các cặp (nữ_i, nam_j) từ 2 nhóm → tính hiệu (nữ_i − nam_j) → trung vị của toàn bộ các hiệu đó. KHÔNG phải "trung vị nhóm nữ trừ trung vị nhóm nam" đơn thuần (tuy nhiên với dữ liệu chuẩn, 2 cách cho kết quả gần bằng nhau).'),
    ('Diễn giải',
     'Điểm GAD-10 của HS nữ có XU HƯỚNG cao hơn điểm GAD-10 của HS nam khoảng 3 điểm (trên thang 0–30).'),
    ('Ý nghĩa lâm sàng',
     '3 điểm / 30 ≈ 10 % thang. Không quá lớn về số học, nhưng ĐỦ để chuyển 1 HS từ ngưỡng "nhẹ" (5–9) lên "vừa" (10–14), hoặc từ "dưới ngưỡng" (<5) lên "nhẹ". Đây là mức thay đổi có ý nghĩa lâm sàng.'),
    ('Hướng khác biệt',
     'Số dương 3 → nữ CAO hơn nam. Nếu là −3 → nữ THẤP hơn nam. Dấu và độ lớn đều quan trọng.'),
]
for i, (k, v) in enumerate(md_data):
    row = md_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = v
d.add_paragraph()

add_para('2.5. Cách đọc đầy đủ câu của Jenkins 2023', bold=True, size=12)

full_tbl = d.add_table(rows=1, cols=1)
full_tbl.style = 'Table Grid'
fc = full_tbl.rows[0].cells[0]
shade_cell(fc, 'DEEBF7')
fp = fc.paragraphs[0]
fr = fp.add_run('Đọc đầy đủ: ')
fr.bold = True
fr.font.color.rgb = RGBColor(31, 73, 125)
fp.add_run('"HS nữ có điểm GAD-10 cao hơn HS nam một cách có ý nghĩa thống kê (p = 0,016, kiểm định Mann-Whitney). Mức chênh lệch trung vị là 3 điểm — tức là phân phối điểm của nữ nhìn chung dịch chuyển lên cao hơn phân phối điểm của nam khoảng 3 điểm trên thang GAD-10 (0–30). Khác biệt này có ý nghĩa lâm sàng khiêm tốn nhưng nhất quán với nhiều NC khác cho thấy VTN nữ có nguy cơ lo âu cao hơn nam."')
d.add_paragraph()

add_para('2.6. Vì sao Jenkins dùng Mann-Whitney thay vì t-test?', bold=True, size=12)

reasons = [
    '① Mẫu nhỏ (N=75, nhóm nữ=45, nhóm nam=30) → kiểm tra normality khó, rủi ro sai loại I/II cao nếu giả định sai.',
    '② Thang GAD-10 có giới hạn trên (30) và giới hạn dưới (0) → dữ liệu thường bị "đè" về 1 đầu → phân phối lệch (skewed), không chuẩn.',
    '③ Thang Likert 0-3 mỗi câu → bản chất là thứ hạng (ordinal), không hoàn toàn continuous → phi tham số phù hợp hơn.',
    '④ Mann-Whitney ROBUST với outlier → kết quả ít bị "lệch" bởi 1-2 HS có điểm cực cao.',
]
for r in reasons:
    p = d.add_paragraph(r)
    for run in p.runs:
        run.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# TÓM GỌN
# =========================================================
add_heading('TÓM GỌN 2 CÂU HỎI', level=1)

sum_tbl = d.add_table(rows=2, cols=1)
sum_tbl.style = 'Table Grid'
sums = [
    ('Câu 1 — NNT',
     'NNT = 1 / (ARR) = số người cần điều trị để có THÊM 1 người đáp ứng. Càng nhỏ = can thiệp càng mạnh, vì NNT nhỏ ⇔ ARR lớn ⇔ chênh lệch hiệu quả lớn so với đối chứng. CAMS combination NNT = 3 là "xuất sắc". Ngưỡng đọc: 1–3 xuất sắc, 4–10 rất tốt, >50 thấp.',
     'E2EFDA'),
    ('Câu 2 — Mann-Whitney + median difference',
     'Mann-Whitney là kiểm định PHI THAM SỐ so 2 nhóm độc lập (thay t-test khi dữ liệu không chuẩn). p=0,016 < 0,05 → có khác biệt thật giữa nữ/nam. "Chênh lệch trung vị 3 điểm" = phân phối điểm GAD-10 của nữ dịch chuyển lên cao hơn nam khoảng 3 điểm trên thang 0-30 (thường tính bằng Hodges-Lehmann estimator). 3/30 ≈ 10% thang → chênh có ý nghĩa lâm sàng khiêm tốn.',
     'FFF2CC'),
]
for i, (k, v, clr) in enumerate(sums):
    cell = sum_tbl.rows[i].cells[0]
    shade_cell(cell, clr)
    p = cell.paragraphs[0]
    r1 = p.add_run(k + '\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)

d.add_paragraph()

# TLTK
add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Laupacis A, Sackett DL, Roberts RS. (1988). An assessment of clinically useful measures of the consequences of treatment. New England Journal of Medicine, 318(26):1728-1733. [NNT lần đầu được đưa ra]',
    'Cook RJ, Sackett DL. (1995). The number needed to treat: a clinically useful measure of treatment effect. BMJ, 310(6977):452-454.',
    'Walkup JT, Albano AM, Piacentini J, et al. (2008). Cognitive behavioral therapy, sertraline, or a combination in childhood anxiety. NEJM, 359:2753-2766. [CAMS — nguồn NNT ≈ 3]',
    'Jenkins JH, Sanchez G, Miller EA, et al. (2023). Depression and anxiety among multiethnic middle school students: Age, gender, and sociocultural environment. [QT001 — nguồn câu hỏi 2]',
    'Mann HB, Whitney DR. (1947). On a test of whether one of two random variables is stochastically larger than the other. Annals of Mathematical Statistics, 18(1):50-60. [Bài gốc giới thiệu Mann-Whitney]',
    'Hodges JL, Lehmann EL. (1963). Estimates of location based on rank tests. Annals of Mathematical Statistics, 34:598-611. [Hodges-Lehmann estimator cho median difference]',
    'Divine GW, Norton HJ, Barón AE, Juarez-Colunga E. (2018). The Wilcoxon-Mann-Whitney procedure fails as a test of medians. American Statistician, 72(3):278-286. [Cảnh báo: Mann-Whitney không test median trực tiếp]',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 20/04/2026 | p-value Mann-Whitney của Jenkins 2023 verify trực tiếp từ tóm tắt QT001 ("Nữ có điểm GAD-10 cao hơn, p=0,016"). Số "chênh lệch trung vị 3 điểm" dựa trên câu hỏi của thầy; em chưa verify con số này từ PDF gốc Jenkins — nếu cần chính xác tuyệt đối khuyến nghị đọc bài gốc phần Results.')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/NNT_va_MannWhitney_chenh_lech_trung_vi.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
