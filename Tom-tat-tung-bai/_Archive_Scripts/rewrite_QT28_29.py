# -*- coding: utf-8 -*-
"""QT28 AJP (2 bảng + 1 hình gốc) + QT29 CBT Meta (3 bảng + 6 hình gốc)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
def make_doc():
    d=Document();s=d.styles['Normal'];s.font.name='Times New Roman';s.font.size=Pt(12)
    s.paragraph_format.space_after=Pt(4);s.paragraph_format.line_spacing=1.5
    for sec in d.sections:sec.top_margin=Cm(2.5);sec.bottom_margin=Cm(2.5);sec.left_margin=Cm(3);sec.right_margin=Cm(2)
    return d
def rb(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=RED
def bl(d,t,bold=False):
    p=d.add_paragraph();r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=BLUE;r.bold=bold
def rh2(d,t):
    h=d.add_heading(t,level=2)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=RED
def bh3(d,t):
    h=d.add_heading(t,level=3)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=BLUE
def shade(c,co):
    s=OxmlElement('w:shd');s.set(qn('w:fill'),co);s.set(qn('w:val'),'clear');c._tc.get_or_add_tcPr().append(s)
def set_w(c,w):
    tcW=c._tc.get_or_add_tcPr();we=OxmlElement('w:tcW');we.set(qn('w:w'),str(int(w*567)));we.set(qn('w:type'),'dxa');tcW.append(we)
def tbl(d,headers,rows,widths):
    t=d.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):set_w(row.cells[ci],widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci];c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:r.font.name='Times New Roman';r.font.size=Pt(10)

# ===== QT28: AJP TREATMENT (2 bảng + 1 hình gốc) =====
d=make_doc()
bl(d,'Tóm tắt bài QT-28',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Phương pháp tiếp cận hiện tại và tương lai trong điều trị rối loạn lo âu ở trẻ em \u00bb (Current and Future Approaches to Pediatric Anxiety Disorder Treatment), do Zugman A và cộng sự (2024), xuất bản trên American Journal of Psychiatry (Q1, IF \u2248 18,0). 12 trang, 2 bảng, 1 hình. Tổng quan toàn diện về CBT, dược lý, và hướng điều trị mới.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Tổng quan tường thuật (narrative review) tổng hợp bằng chứng về các phương pháp điều trị rối loạn lo âu ở trẻ em và thanh thiếu niên. Nói cách khác, đây là bài phân tích tổng hợp từ tạp chí tâm thần hàng đầu thế giới, cung cấp bức tranh toàn cảnh về "hiện tại và tương lai" của điều trị lo âu trẻ em.',bold=True)

bl(d,'Tổng quan tài liệu cho thấy CBT và SSRI là hai phương pháp đã được xác lập (established treatments). Tuy nhiên, khoảng 30\u201340% trẻ không đáp ứng với liệu pháp đầu tiên, đòi hỏi các hướng tiếp cận mới. Đặc biệt, bằng chứng từ LMIC (nước thu nhập thấp-trung bình) rất hạn chế \u2014 0 RCT từ Việt Nam (Zhameden 2025).',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 1 gốc). Hiệu quả các phương pháp điều trị lo âu trẻ em:', bold=True)
tbl(d,['Phương pháp','Hiệu quả','Bằng chứng','Ghi chú'],
    [['CBT cá nhân','Phục hồi 47\u201366%\nĐáp ứng 57\u201360%','RCTs nhiều','Bằng chứng mạnh nhất'],
     ['SSRI (thuốc chống trầm cảm)','Hiệu quả vừa phải','RCTs nhiều','Kết hợp CBT tốt hơn'],
     ['CBT + SSRI kết hợp','Tốt hơn đơn trị','CAMS trial','Chuẩn vàng hiện tại'],
     ['CBT nhóm','Hiệu quả tương đương CBT cá nhân','RCTs','Chi phí thấp hơn'],
     ['CBT qua internet (iCBT)','Triển vọng','RCTs sơ bộ','Tăng tiếp cận'],
     ['Kích thích thần kinh (TMS, tDCS)','Thử nghiệm','Ít RCT','Chưa đủ bằng chứng'],
     ['Ứng dụng kỹ thuật số','Triển vọng','Pilot studies','Cần thêm NC']],
    widths=[3.5,3.0,2.5,3.5])
d.add_paragraph()

bl(d,'Bảng 2. So sánh bằng chứng can thiệp lo âu từ nhiều nguồn:', bold=True)
tbl(d,['Nguồn','Phương pháp','Phát hiện chính','Hạn chế'],
    [['AJP 2024 (bài này)','Tổng quan','CBT 47\u201366% phục hồi','Chủ yếu phương Tây'],
     ['Zhameden 2025','6 RCTs LMIC','CBT 3/4 trầm cảm, 1/4 lo âu','GRADE rất thấp, 0 VN'],
     ['BMC 2025 (QT29)','30 RCTs network meta','CBT cá nhân tốt nhất','Ít LMIC'],
     ['Wen 2020','Quan sát','Hỗ trợ SKTT trường OR=0,562','Cắt ngang'],
     ['Hoàng Trung Học VN','Hồi quy','Gia đình Beta=0,272','Không can thiệp']],
    widths=[3.0,3.0,4.0,3.0])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : CBT được xác nhận hiệu quả nhất (47\u201366% phục hồi) bởi cả AJP 2024, BMC network meta 2025 (30 RCTs), và Zhameden 2025. Tuy nhiên, CBT chỉ hiệu quả 1/4 cho lo âu tại LMIC (Zhameden) \u2014 gợi ý cần điều chỉnh cho bối cảnh VN. Hướng mới (iCBT, ứng dụng kỹ thuật số) đặc biệt phù hợp cho VN nơi tiếp cận dịch vụ chỉ 8,4%.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*CBT + SSRI = chuẩn vàng.* CAMS trial (lớn nhất) cho thấy kết hợp hiệu quả hơn đơn trị. Tuy nhiên, SSRI có tác dụng phụ ở trẻ \u2014 cần cân nhắc.')
bl(d,'*30\u201340% không đáp ứng.* Đây là khoảng trống lớn \u2014 cần hướng mới cho nhóm kháng trị.')
bl(d,'*iCBT và ứng dụng kỹ thuật số.* Đặc biệt phù hợp cho VN: tăng tiếp cận ở vùng thiếu chuyên gia, chi phí thấp, mở rộng quy mô nhanh. Phù hợp với phát hiện Nature 2025 về VTN và công nghệ.')

rb(d,'Kết luận')
bl(d,'Dữ liệu tổng hợp từ AJP (IF = 18), cho thấy CBT là phương pháp hiệu quả nhất (47\u201366% phục hồi) nhưng 30\u201340% không đáp ứng, gợi ý rằng cần phát triển hướng mới: iCBT, ứng dụng kỹ thuật số, và kích thích thần kinh. Tại VN nơi 0 RCT và 8,4% tiếp cận dịch vụ, iCBT có thể là giải pháp khả thi nhất để mở rộng can thiệp.',bold=True)

rh2(d,'Phản biện')
bl(d,'AJP Q1 IF = 18 \u2014 tạp chí tâm thần hàng đầu thế giới. Tổng quan toàn diện, cập nhật nhất. Tuy nhiên, tổng quan (review) không có dữ liệu gốc mới. Chủ yếu bằng chứng từ phương Tây (Mỹ, Anh, Úc) \u2014 thiếu dữ liệu LMIC/châu Á. Hình 1 gốc về cơ chế thần kinh cần chèn từ PDF.')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'RCT CBT nhóm tại trường THCS/THPT VN \u2014 RCT đầu tiên. Thử nghiệm iCBT tiếng Việt cho VTN. Đánh giá hiệu quả CBT kết hợp can thiệp gia đình (dựa trên Pham 2024: chăm sóc cảm xúc beta = \u20130,40).')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. AJP Q1 IF = 18, tổng quan toàn diện nhất về điều trị lo âu trẻ em.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT28_AJP_Treatment_2024.docx')
sys.stderr.write('QT28 OK\n')

# ===== QT29: CBT NETWORK META (3 bảng + 6 hình gốc) =====
d=make_doc()
bl(d,'Tóm tắt bài QT-29',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Hiệu quả các loại can thiệp khác nhau cho rối loạn lo âu ở trẻ em và thanh thiếu niên: Tổng quan hệ thống và phân tích tổng hợp mạng Bayesian \u00bb (Effects of Different Interventions on Anxiety Disorders in Children and Adolescents: A Systematic Review and Bayesian Network Meta-analysis), xuất bản trên BMC Psychiatry, 2025 (Q1, Open Access). 30 RCTs, 1.711 trẻ tham gia. 14 trang, 3 bảng, 6 hình.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Tổng quan hệ thống và phân tích tổng hợp mạng Bayesian (Bayesian network meta-analysis) \u2014 phương pháp nâng cao nhất cho phép so sánh ĐỒNG THỜI nhiều loại can thiệp mà không cần RCT head-to-head cho mọi cặp. Nói cách khác, thay vì chỉ so sánh CBT vs placebo, phương pháp này xếp hạng tất cả các loại can thiệp dựa trên mạng lưới bằng chứng.',bold=True)

bl(d,'Tìm kiếm trong PubMed, PsychINFO, Web of Science đến 24/01/2025. Tiêu chí: RCT, trẻ em/VTN có rối loạn lo âu, so sánh ít nhất 2 can thiệp. Đánh giá nguy cơ thiên lệch bằng Cochrane RoB 2.0.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 1 gốc). Đặc điểm 30 RCTs:', bold=True)
tbl(d,['Đặc điểm','Giá trị'],
    [['Số RCTs','30'],
     ['Tổng số trẻ tham gia','1.711'],
     ['Loại can thiệp so sánh','CBT cá nhân, CBT nhóm, thuốc (SSRI), kết hợp, tâm lý giáo dục, thư giãn, chờ đợi'],
     ['Rối loạn lo âu đánh giá','GAD, SAD, lo âu xã hội, lo âu phân ly, ám ảnh cụ thể'],
     ['Phương pháp phân tích','Bayesian network meta-analysis (NMA)']],
    widths=[4.0,8.5])
d.add_paragraph()

bl(d,'Bảng 2 (từ Table 2+3 gốc). Xếp hạng hiệu quả các can thiệp:', bold=True)
tbl(d,['Xếp hạng','Can thiệp','Hiệu quả (SMD/OR)','Ghi chú'],
    [['1','CBT cá nhân','Hiệu quả nhất','Bằng chứng mạnh'],
     ['2\u20133','CBT nhóm','Tương đương cá nhân','Chi phí thấp hơn'],
     ['3\u20134','SSRI (thuốc)','Hiệu quả','Tác dụng phụ ở trẻ'],
     ['4\u20135','CBT + SSRI','Tốt nhưng ít RCT','Cần thêm NC'],
     ['5+','Tâm lý giáo dục','Hiệu quả thấp hơn','An toàn'],
     ['','Thư giãn','Hiệu quả hạn chế',''],
     ['','Chờ đợi (waitlist)','Tham chiếu','Không can thiệp']],
    widths=[1.5,3.0,3.5,4.5])
d.add_paragraph()

bl(d,'Bảng 3. So sánh với các tổng quan can thiệp khác:', bold=True)
tbl(d,['Nguồn','Phạm vi','Phát hiện','So với bài này'],
    [['BMC 2025 (bài này)','30 RCTs, 1.711 trẻ','CBT cá nhân tốt nhất','NMA Bayesian \u2014 nâng cao nhất'],
     ['AJP 2024 (QT28)','Tổng quan tường thuật','CBT 47\u201366% phục hồi','Không xếp hạng'],
     ['Zhameden 2025','6 RCTs LMIC','CBT 3/4 dep, 1/4 anx','GRADE rất thấp, 0 VN'],
     ['Wen 2020','Quan sát nông thôn TQ','SKTT trường OR = 0,562','Không can thiệp']],
    widths=[3.0,3.0,3.5,3.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : Ba tổng quan can thiệp (AJP 2024, BMC 2025, Zhameden 2025) đều xác nhận CBT hiệu quả nhất. Tuy nhiên, Zhameden cảnh báo CBT chỉ hiệu quả 1/4 cho lo âu tại LMIC \u2014 gợi ý bằng chứng từ phương Tây chưa chắc áp dụng được cho VN. 30 RCTs trong BMC 2025 chủ yếu từ phương Tây \u2014 khoảng trống LMIC/châu Á rất lớn.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*CBT cá nhân xếp hạng 1.* Network meta Bayesian xác nhận CBT cá nhân hiệu quả nhất trong 7+ loại can thiệp \u2014 phù hợp AJP 2024.')
bl(d,'*CBT nhóm \u2014 lựa chọn thực tế cho VN.* Hiệu quả tương đương CBT cá nhân nhưng chi phí thấp hơn, phù hợp bối cảnh trường học VN (có thể triển khai cho lớp 20\u201330 HS).')
bl(d,'*0 RCT từ VN.* Trong 30 RCTs của NMA, không có bài nào từ Việt Nam hay ASEAN \u2014 xác nhận khoảng trống Zhameden 2025 đã chỉ ra.')
bl(d,'*6 hình gốc.* Bài gốc có 6 hình (forest plots, network plots, funnel plots) \u2014 cần chèn trong bản dịch đầy đủ.')

rb(d,'Kết luận')
bl(d,'Dữ liệu từ phân tích tổng hợp mạng Bayesian trên 30 RCTs (1.711 trẻ), cho thấy CBT cá nhân hiệu quả nhất, CBT nhóm xếp thứ 2\u20133 (tương đương nhưng chi phí thấp hơn), gợi ý rằng CBT nhóm tại trường học là lựa chọn khả thi nhất cho RCT đầu tiên tại Việt Nam. Khoảng trống 0 RCT từ VN/ASEAN nhấn mạnh nhu cầu cấp thiết cho nghiên cứu can thiệp tại khu vực này.',bold=True)

rh2(d,'Phản biện')
bl(d,'BMC Psychiatry Q1, Open Access. NMA Bayesian \u2014 phương pháp nâng cao nhất, Cochrane RoB 2.0. 30 RCTs là đủ cho NMA. Tuy nhiên, 1.711 trẻ \u2014 cỡ mẫu tổng không lớn. Chủ yếu RCTs phương Tây \u2014 thiên lệch bối cảnh. Cần đọc chi tiết forest plots (6 hình) để đánh giá heterogeneity. Có thể có publication bias (kiểm tra funnel plot).')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'RCT CBT nhóm tại trường THCS/THPT VN \u2014 đóng góp vào NMA toàn cầu. Điều chỉnh CBT cho bối cảnh văn hóa VN (gia đình, tập thể). Thêm dữ liệu LMIC/châu Á vào NMA. Đánh giá iCBT vs CBT nhóm trực tiếp.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50 Cao. BMC Q1, 30 RCTs, NMA Bayesian nâng cao, Cochrane RoB 2.0. Nhưng chủ yếu phương Tây.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT29_CBT_NetworkMeta_2025.docx')
sys.stderr.write('QT29 OK\n')
