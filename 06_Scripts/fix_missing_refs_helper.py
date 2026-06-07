"""SKELETON — helper to append a 'Tham khảo' section to summary docx files.

STATUS: TEMPLATE ONLY. DO NOT RUN BLINDLY.
Author workflow REQUIRES per-file:
  1) Read PDF goc trong `02_Papers-goc/` -> extract authoritative reference list
  2) Verify DOI / PMID qua Crossref + PubMed (see feedback_verify_tung_entry_truoc_khi_gui.md)
  3) Only then append. NEVER fabricate DOI from memory.

Usage (after curating per-file refs):
    python fix_missing_refs_helper.py --file QT001_Jenkins_et_al_2023_USA.docx --refs refs_QT001.json

Where refs_QT001.json is human-curated, e.g.:
  {
    "primary_doi": "10.1037/xxx",
    "primary_pmid": "PMID:12345",
    "primary_url": "https://doi.org/10.1037/xxx",
    "extra_refs": [
      {"apa": "Smith J, et al. (2023). Title. Journal vol(iss):pp. DOI", "doi": "10.xxx/yyy", "verified_by": "Crossref"}
    ],
    "notes": "Cross-checked with PDF goc page 12-14, all 7 refs match."
  }
"""
from __future__ import annotations

import argparse
import json
from pathlib import Path
from datetime import date

from docx import Document
from docx.shared import Pt

TOM_TAT_DIR = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au/Tom-tat-tung-bai")


def append_reference_section(docx_path: Path, refs_payload: dict, save_as: Path | None = None) -> Path:
    """Append a 'Tham khao' (APA 7) section to the docx.

    refs_payload format documented in module docstring.
    By default writes to <name>_FIXED_<DDMMYYYY>.docx so the original is untouched.
    """
    doc = Document(str(docx_path))

    # Heading
    h = doc.add_paragraph()
    run = h.add_run("Tham khảo (APA 7)")
    run.bold = True
    run.font.size = Pt(13)

    # Primary refs (DOI / PMID / URL of the source paper)
    primary_lines = []
    if refs_payload.get("primary_doi"):
        primary_lines.append(f"DOI: {refs_payload['primary_doi']}")
    if refs_payload.get("primary_pmid"):
        primary_lines.append(refs_payload["primary_pmid"])
    if refs_payload.get("primary_url"):
        primary_lines.append(f"URL: {refs_payload['primary_url']}")
    for line in primary_lines:
        doc.add_paragraph(f"• {line}")

    # Extra references (curated)
    for r in refs_payload.get("extra_refs", []) or []:
        p = doc.add_paragraph(f"• {r['apa']}")
        if r.get("doi"):
            p.add_run(f"  [DOI: {r['doi']}]")

    # Provenance footer
    notes = refs_payload.get("notes", "")
    doc.add_paragraph(
        f"Nguồn kiểm chứng: PDF gốc + Crossref/PubMed. "
        f"Ngày cập nhật: {date.today().isoformat()}. {notes}"
    )

    target = save_as or docx_path.with_name(
        docx_path.stem + f"_FIXED_{date.today().strftime('%d%m%Y')}.docx"
    )
    doc.save(str(target))
    return target


def parse_args():
    ap = argparse.ArgumentParser()
    ap.add_argument("--file", required=True, help="docx filename inside Tom-tat-tung-bai/")
    ap.add_argument("--refs", required=True, help="JSON file with curated refs payload")
    ap.add_argument("--out", default=None, help="optional output path")
    return ap.parse_args()


def main():
    args = parse_args()
    src = TOM_TAT_DIR / args.file
    if not src.exists():
        raise SystemExit(f"Not found: {src}")
    refs = json.loads(Path(args.refs).read_text(encoding="utf-8"))
    out = append_reference_section(src, refs, Path(args.out) if args.out else None)
    print(f"Wrote: {out}")


if __name__ == "__main__":
    # Guard: prevent accidental batch run
    raise SystemExit(
        "SKELETON only. Curate refs JSON per-file then invoke explicitly:\n"
        "  python fix_missing_refs_helper.py --file <name>.docx --refs <name>.json"
    )
    # main()  # noqa: enable after curation
