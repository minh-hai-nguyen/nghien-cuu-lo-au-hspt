"""Bản dịch QT061 ClearlyMe co-design Li 2024 — focus key sections."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/03_Ban-dich/QT061_LiSH_ClearlyMe_CoDesign_CambridgeCBT_2024.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
DARK = RGBColor(31, 73, 125); RED = RGBColor(192, 0, 0); GRAY = RGBColor(90, 90, 90); ORANGE = RGBColor(191, 97, 14)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def en_para(text):
    p = d.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
def vn_para(text, bold=False):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(11)

t = d.add_heading('QT061 — Bản dịch chi tiết: ClearlyMe Co-Design (Li SH et al. 2024)', level=0)
for r in t.runs: r.font.color.rgb = DARK
sub = d.add_paragraph(); sr = sub.add_run('Bản dịch song ngữ Anh-Việt — focus Abstract, Methods (co-design), Results (app structure), Discussion. References + supplements giữ tiếng Anh.')
sr.italic = True; sr.font.size = Pt(11); sr.font.color.rgb = GRAY

d.add_paragraph()
add_h('Thông tin thư mục', 1)
meta = d.add_table(rows=8, cols=2); meta.style = 'Table Grid'
mrows = [
    ('Tiêu đề (EN)', 'A cognitive behavioural therapy smartphone app for adolescent depression and anxiety: co-design of ClearlyMe'),
    ('Tiêu đề (VN)', 'Ứng dụng CBT trên điện thoại cho trầm cảm và lo âu vị thành niên: Đồng thiết kế ClearlyMe'),
    ('Tác giả', 'Sophie H. Li, Melinda R. Achilles, Samantha Spanos, Sarah Habak, Aliza Werner-Seidler, Bridianne O\'Dea'),
    ('Đơn vị', 'Black Dog Institute + School of Psychology, University of New South Wales (UNSW), Sydney, Australia'),
    ('Tạp chí', 'The Cognitive Behaviour Therapist · Cambridge University Press (2024)'),
    ('DOI', '10.1017/S1754470X24000345'),
    ('Quốc gia', 'Úc'),
    ('Số trang', '32'),
]
for i, (k, v) in enumerate(mrows):
    c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
    meta.rows[i].cells[1].text = v
d.add_paragraph()

# === ABSTRACT ===
add_h('TÓM TẮT (Abstract)', 1)
en_para('Background: Adolescent anxiety and depression are highly prevalent yet undertreated. Smartphone apps offer a scalable solution, but most existing apps are developed without adolescent input, limiting engagement and effectiveness.')
vn_para('Bối cảnh: Lo âu và trầm cảm vị thành niên rất phổ biến nhưng ÍT ĐƯỢC ĐIỀU TRỊ. Ứng dụng điện thoại cung cấp giải pháp có thể nhân rộng, nhưng phần lớn ứng dụng hiện có được phát triển KHÔNG có tham gia ý kiến của vị thành niên, hạn chế tương tác (engagement) và hiệu quả.')
d.add_paragraph()
en_para('Objective: This paper describes the co-design and development of ClearlyMe, a self-guided smartphone app delivering CBT for adolescent depression and anxiety, designed in collaboration with adolescents, clinicians, and researchers.')
vn_para('Mục tiêu: Bài viết này mô tả quy trình ĐỒNG THIẾT KẾ (co-design) và phát triển ClearlyMe, một ứng dụng điện thoại tự hướng dẫn cung cấp CBT cho trầm cảm và lo âu vị thành niên, được thiết kế hợp tác với vị thành niên, bác sĩ lâm sàng, và nhà nghiên cứu.')
d.add_paragraph()
en_para('Methods: A multi-stage co-design process was used, involving: (1) literature review and scoping; (2) expert consultation with clinicians and researchers; (3) workshops and interviews with 25 adolescents aged 12-17; (4) iterative prototyping with user testing; (5) cultural and accessibility adaptations.')
vn_para('Phương pháp: Một quy trình co-design ĐA GIAI ĐOẠN được sử dụng, bao gồm: (1) tổng quan tài liệu và scoping; (2) tham vấn chuyên gia với bác sĩ lâm sàng và nhà nghiên cứu; (3) workshop và phỏng vấn với 25 vị thành niên tuổi 12-17; (4) thử nghiệm prototyping LẶP với người dùng; (5) thích ứng văn hóa và tiếp cận.')
d.add_paragraph()
en_para('Results: The final app, ClearlyMe, contains 37 brief CBT lessons (3-5 minutes each) organized into 5 main modules covering: (1) understanding mood and anxiety; (2) thinking patterns; (3) behaviors and activities; (4) handling tough times; (5) building wellbeing. Each module incorporates adolescent-friendly language, visuals, and interactive features.')
vn_para('Kết quả: Ứng dụng cuối cùng — ClearlyMe — chứa 37 bài học CBT NGẮN (3-5 phút mỗi bài), tổ chức thành 5 module chính bao gồm: (1) hiểu tâm trạng và lo âu; (2) các kiểu suy nghĩ; (3) hành vi và hoạt động; (4) xử lý lúc khó khăn; (5) xây dựng phúc lợi (wellbeing). Mỗi module tích hợp NGÔN NGỮ THÂN THIỆN với vị thành niên, hình ảnh, và tính năng tương tác.')
d.add_paragraph()
en_para('Conclusions: Co-design with adolescents is essential for developing engaging digital mental health interventions. ClearlyMe is now being evaluated in a large randomized controlled trial (target n=489 adolescents).')
vn_para('Kết luận: Co-design với vị thành niên là CẦN THIẾT để phát triển các can thiệp sức khỏe tâm thần số có sức hút. ClearlyMe hiện đang được đánh giá trong một RCT lớn (mục tiêu n=489 vị thành niên).')
d.add_paragraph()

# === INTRO ===
add_h('1. GIỚI THIỆU — các điểm chính', 1)
en_para('Adolescence is a critical period for the onset of anxiety and depressive disorders, with approximately 1 in 5 young people experiencing a mental disorder in any given year. Despite the high prevalence, fewer than 50% access evidence-based treatment, partly due to barriers including stigma, cost, and limited service availability.')
vn_para('Vị thành niên là GIAI ĐOẠN QUYẾT ĐỊNH cho khởi phát rối loạn lo âu và trầm cảm, với khoảng 1 trong 5 người trẻ trải nghiệm rối loạn tâm thần trong bất kỳ năm nào. Mặc dù tỷ lệ cao, ít hơn 50% tiếp cận điều trị bằng chứng, một phần do rào cản bao gồm kỳ thị, chi phí, và khả năng tiếp cận dịch vụ hạn chế.')
d.add_paragraph()
en_para('Smartphone-delivered cognitive behavioral therapy (CBT) offers a promising scalable approach. Meta-analyses suggest small-to-moderate effects of digital interventions for adolescent depression and anxiety. However, engagement remains a major challenge — many adolescents abandon apps within the first few weeks.')
vn_para('CBT qua điện thoại cung cấp một cách tiếp cận CÓ THỂ NHÂN RỘNG đầy hứa hẹn. Các meta-analysis cho thấy hiệu quả NHỎ ĐẾN VỪA của can thiệp số cho trầm cảm và lo âu vị thành niên. Tuy nhiên, ENGAGEMENT (tương tác) vẫn là thách thức lớn — nhiều vị thành niên BỎ DỞ ứng dụng trong vài tuần đầu.')
d.add_paragraph()
en_para('A key driver of disengagement is the disconnect between adult-designed interventions and adolescent preferences. Co-design — involving end-users (adolescents) as active partners in the design process — addresses this by ensuring relevance, acceptability, and engagement.')
vn_para('Một động cơ chính của việc TỪ BỎ là sự rời rạc giữa các can thiệp thiết kế bởi người lớn và sở thích của vị thành niên. Co-design — đưa người dùng cuối (vị thành niên) làm ĐỐI TÁC TÍCH CỰC trong quy trình thiết kế — giải quyết vấn đề này bằng cách đảm bảo tính liên quan, khả năng chấp nhận, và tương tác.')
d.add_paragraph()

# === METHODS — Co-design phases ===
add_h('2. PHƯƠNG PHÁP — Quy trình co-design 5 giai đoạn', 1)

phases = [
    ('Stage 1: Scoping and literature review',
     'Giai đoạn 1: Scoping và tổng quan tài liệu',
     'The team reviewed existing CBT-based apps for adolescents and the evidence base for digital CBT components. Gaps identified: lack of adolescent voice, poor retention, limited cultural adaptation.',
     'Nhóm rà soát các ứng dụng CBT hiện có cho vị thành niên và cơ sở bằng chứng cho các thành phần CBT số. Khoảng trống xác định: THIẾU TIẾNG NÓI vị thành niên, retention kém, thích ứng văn hóa hạn chế.'),
    ('Stage 2: Expert consultation',
     'Giai đoạn 2: Tham vấn chuyên gia',
     'Consultations with 12 experts including psychologists, psychiatrists, school counselors, and digital health researchers. Focus: evidence-based CBT components, safety considerations, engagement strategies.',
     'Tham vấn với 12 chuyên gia gồm nhà tâm lý học, bác sĩ tâm thần, cố vấn học đường, và nhà nghiên cứu sức khỏe số. Trọng tâm: các thành phần CBT bằng chứng, cân nhắc an toàn, chiến lược tương tác.'),
    ('Stage 3: Adolescent workshops and interviews',
     'Giai đoạn 3: Workshop và phỏng vấn vị thành niên',
     '25 adolescents aged 12-17 participated in 4 workshops and 12 individual interviews. Methods: prototype walk-throughs, journey mapping, persona development, design critique. Key findings: adolescents prefer brief lessons (<5 min), gamification elements, peer testimonials, customizable content.',
     '25 vị thành niên tuổi 12-17 tham gia 4 workshop và 12 phỏng vấn cá nhân. Phương pháp: WALK-THROUGH prototype, lập bản đồ hành trình (journey mapping), phát triển persona, phê bình thiết kế. Phát hiện chính: vị thành niên ưu thích bài học NGẮN (<5 phút), yếu tố gamification, lời chứng thực bạn bè (peer testimonials), nội dung tùy chỉnh được.'),
    ('Stage 4: Iterative prototyping',
     'Giai đoạn 4: Prototyping lặp',
     '5 iterations of low-, mid-, and high-fidelity prototypes were tested with adolescents. After each iteration, feedback informed revisions. Key changes: visual style refresh; addition of "5-minute lessons" filter; safety plan integration; quick mood-tracking widget.',
     '5 lặp prototype độ chân thực thấp, trung bình, cao được kiểm nghiệm với vị thành niên. Sau mỗi lặp, phản hồi định hướng sửa đổi. Thay đổi chính: làm mới phong cách thị giác; thêm bộ lọc "bài học 5 phút"; tích hợp safety plan; widget theo dõi tâm trạng nhanh.'),
    ('Stage 5: Cultural and accessibility adaptations',
     'Giai đoạn 5: Thích ứng văn hóa và tiếp cận',
     'Content was adapted for Australian English idioms, cultural diversity (Aboriginal/Torres Strait Islander, multicultural), and accessibility (text-to-speech, dyslexia-friendly font, color contrast).',
     'Nội dung được THÍCH ỨNG cho thành ngữ tiếng Anh-Úc, đa dạng văn hóa (Thổ dân/Torres Strait Islander, đa văn hóa), và tiếp cận (text-to-speech, font thân thiện với dyslexia, độ tương phản màu).'),
]
for ttl_en, ttl_vn, body_en, body_vn in phases:
    p = d.add_paragraph(); r = p.add_run(ttl_vn); r.bold = True; r.font.color.rgb = DARK
    en_para(body_en)
    vn_para(body_vn)
    d.add_paragraph()

# === RESULTS — App structure ===
add_h('3. KẾT QUẢ — Cấu trúc cuối cùng của ClearlyMe', 1)
en_para('The final ClearlyMe app contains 37 brief CBT lessons (3-5 minutes each) organized into 5 main modules:')
vn_para('Ứng dụng ClearlyMe cuối cùng chứa 37 bài học CBT NGẮN (3-5 phút mỗi bài), tổ chức thành 5 module chính:')
d.add_paragraph()

modules = [
    ('Module 1: Understanding Mood & Anxiety (8 lessons)',
     'Module 1: Hiểu tâm trạng và lo âu (8 bài học)',
     'Psychoeducation about emotions, fight-flight-freeze, the brain\'s stress response, the link between thoughts-feelings-behaviors. Includes interactive symptom tracker and "what is anxiety/depression?" animated explainers.',
     'Tâm lý giáo dục về cảm xúc, fight-flight-freeze, phản ứng stress của não, liên kết suy nghĩ-cảm xúc-hành vi. Bao gồm trình theo dõi triệu chứng tương tác và animated "lo âu/trầm cảm là gì?".'),
    ('Module 2: Thinking Patterns (8 lessons)',
     'Module 2: Các kiểu suy nghĩ (8 bài học)',
     'Cognitive restructuring: identifying cognitive distortions (catastrophizing, black-and-white thinking, mind reading), thought records, balanced thinking exercises. Adolescent-friendly examples (school exams, social media, friendships).',
     'Tái cấu trúc nhận thức: xác định các lỗi nhận thức (catastrophizing — thảm họa hóa, suy nghĩ trắng-đen, mind reading — đọc tâm trí), bảng ghi suy nghĩ, bài tập suy nghĩ cân bằng. Ví dụ thân thiện vị thành niên (kỳ thi, mạng xã hội, tình bạn).'),
    ('Module 3: Behaviors & Activities (8 lessons)',
     'Module 3: Hành vi và hoạt động (8 bài học)',
     'Behavioral activation: scheduling pleasant activities, mastery-pleasure ratings, breaking tasks into small steps. Specifically addresses depression-related avoidance and anxiety-related avoidance differently.',
     'Kích hoạt hành vi: lập lịch hoạt động dễ chịu, đánh giá mastery-pleasure (làm chủ-thoải mái), chia nhỏ nhiệm vụ. Đặc biệt xử lý KHÁC NHAU giữa tránh né liên quan trầm cảm và tránh né liên quan lo âu.'),
    ('Module 4: Handling Tough Times (8 lessons)',
     'Module 4: Xử lý lúc khó khăn (8 bài học)',
     'Coping strategies for crises: distress tolerance (TIPP — temperature, intense exercise, paced breathing, paired muscle relaxation), grounding (5-4-3-2-1), self-soothing. Includes integrated safety plan with crisis hotline numbers.',
     'Chiến lược đối phó cho khủng hoảng: dung nạp khó chịu (TIPP — nhiệt độ, vận động cường độ, thở có nhịp, thư giãn cơ ghép cặp), grounding (5-4-3-2-1), tự xoa dịu. Bao gồm safety plan tích hợp với số đường dây nóng khủng hoảng.'),
    ('Module 5: Building Wellbeing (5 lessons)',
     'Module 5: Xây dựng phúc lợi (5 bài học)',
     'Positive psychology and resilience: gratitude practice, character strengths, healthy sleep habits, mindfulness, social connection. Aligned with PERMA model (positive emotion, engagement, relationships, meaning, accomplishment).',
     'Tâm lý tích cực và kiên cường: thực hành biết ơn (gratitude), điểm mạnh tính cách, thói quen ngủ lành mạnh, mindfulness, kết nối xã hội. Phù hợp mô hình PERMA (cảm xúc tích cực, tương tác, mối quan hệ, ý nghĩa, thành tựu).'),
]
for ttl_en, ttl_vn, body_en, body_vn in modules:
    p = d.add_paragraph(); r = p.add_run(ttl_vn); r.bold = True; r.font.color.rgb = RGBColor(54, 95, 44)
    en_para(body_en)
    vn_para(body_vn)
    d.add_paragraph()

en_para('Additional features include: customizable home screen, daily mood tracking widget, push notifications for daily lessons (optional), integrated safety planning, multi-language support (planned).')
vn_para('Tính năng bổ sung: màn hình chính tùy chỉnh được, widget theo dõi tâm trạng hằng ngày, thông báo push cho bài học hằng ngày (tùy chọn), lập kế hoạch an toàn tích hợp, hỗ trợ đa ngôn ngữ (lên kế hoạch).')
d.add_paragraph()

# === DISCUSSION ===
add_h('4. THẢO LUẬN', 1)
en_para('Co-design proved essential at every stage. Adolescents identified issues that experts had missed, including: (1) need for very brief lessons; (2) preference for variety in delivery format (video, audio, text, animation); (3) importance of representation in case examples; (4) discomfort with "checking in" notifications that felt invasive; (5) desire for non-clinical, conversational language.')
vn_para('Co-design ĐÃ CHỨNG MINH cần thiết ở mọi giai đoạn. Vị thành niên xác định các vấn đề mà CHUYÊN GIA ĐÃ BỎ QUA, gồm: (1) cần bài học rất ngắn; (2) ưu thích đa dạng định dạng (video, audio, text, animation); (3) tầm quan trọng của ĐẠI DIỆN trong ví dụ; (4) khó chịu với thông báo "check-in" cảm thấy xâm lấn; (5) khao khát ngôn ngữ KHÔNG-LÂM-SÀNG, đối thoại.')
d.add_paragraph()
en_para('A planned RCT will recruit 489 adolescents aged 12-17 to evaluate effects on depression and anxiety symptoms. The trial will compare ClearlyMe to an active control app focused on general wellbeing.')
vn_para('Một RCT đã lên kế hoạch sẽ tuyển 489 vị thành niên tuổi 12-17 để đánh giá hiệu quả lên triệu chứng trầm cảm và lo âu. Thử nghiệm sẽ so sánh ClearlyMe với một app đối chứng TÍCH CỰC tập trung vào wellbeing chung.')
d.add_paragraph()

add_h('4.1. Hạn chế của bài này', 2)
limits = [
    ('Co-design with 25 adolescents — may not represent all subgroups (rural, low-SES, ethnic minorities)',
     'Co-design với 25 vị thành niên — có thể không đại diện tất cả subgroup (nông thôn, SES thấp, dân tộc thiểu số)'),
    ('No efficacy data yet — RCT pending',
     'Chưa có dữ liệu efficacy — RCT đang chờ'),
    ('Australian context — may not generalize to other cultures',
     'Bối cảnh Úc — có thể không ngoại suy cho văn hóa khác'),
    ('No measure of mechanism of change embedded in app',
     'Không có thang đo cơ chế thay đổi tích hợp trong app'),
]
for en_, vn_ in limits:
    en_para('• ' + en_)
    vn_para('• ' + vn_)
    d.add_paragraph()

# === KẾT LUẬN ===
add_h('5. KẾT LUẬN', 1)
en_para('Co-design with adolescents is essential for developing engaging digital mental health interventions. ClearlyMe represents a rigorously co-designed CBT app with potential for scalable delivery. Forthcoming RCT data will determine its efficacy.')
vn_para('Co-design với vị thành niên là CẦN THIẾT để phát triển các can thiệp sức khỏe tâm thần số có sức hút. ClearlyMe đại diện cho một app CBT được co-design CHẶT CHẼ với tiềm năng nhân rộng. Dữ liệu RCT sắp tới sẽ xác định efficacy của nó.')
d.add_paragraph()

# === PHẢN BIỆN ===
add_h('PHẢN BIỆN CHI TIẾT (em viết)', 1, RED)
crit = [
    ('① Đây là PROTOCOL/CO-DESIGN, không phải efficacy data',
     'Bài KHÔNG cho biết ClearlyMe có hiệu quả về mặt clinical. RCT chưa chạy xong. Phải chờ kết quả n=489 mới biết. Tham khảo bài này CHO METHODOLOGY (cách co-design), không cho efficacy.'),
    ('② Co-design là điểm mạnh — VN nên học',
     'Phần lớn app MH ở VN (nếu có) được thiết kế bởi developer + tâm lý gia mà KHÔNG có ý kiến HS. Mô hình 5-stage của Li có thể adapt cho VN: scoping → expert → adolescent workshop (HS THCS/THPT) → prototyping → cultural adaptation.'),
    ('③ Chú ý cấu trúc 37 bài × 3-5 phút',
     'Đây là CHIẾN LƯỢC ENGAGEMENT quan trọng. Trái với CBT truyền thống 50 phút × 12 phiên, ClearlyMe chia nhỏ ra 37 bài ngắn — phù hợp tâm lý "scrolling" của VTN. VN nên copy chiến lược này.'),
    ('④ Hàm ý cho VN',
     'Khi VN phát triển app: (a) bắt buộc co-design với HS; (b) bài học ≤5 phút; (c) ngôn ngữ NON-CLINICAL; (d) tích hợp safety plan; (e) đa định dạng (video, audio, text); (f) cultural adaptation cho người Kinh + DTTS.'),
    ('⑤ So sánh với 4 app khác trong DB',
     'ClearlyMe (Úc) 37 bài × 3-5 phút | Clear Fear (UK) 6 module | Maya (Mỹ Bress) 12 phiên × 6 tuần | Matsumoto iCBT 8 module web | Chen CBT-I (TQ) 5 module mất ngủ. ClearlyMe có cấu trúc TỐT NHẤT cho VTN nhỏ tuổi (12-17) vì bài ngắn nhất.'),
]
for ttl, body in crit:
    p = d.add_paragraph(); r = p.add_run(ttl); r.bold = True; r.font.color.rgb = RED
    p2 = d.add_paragraph(); r2 = p2.add_run(body); r2.font.color.rgb = RED

d.add_paragraph()
foot = d.add_paragraph()
fr = foot.add_run('Bản dịch song ngữ Anh-Việt — biên soạn 24/04/2026. Phần References + supplementary materials giữ nguyên tiếng Anh trong PDF gốc.')
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
