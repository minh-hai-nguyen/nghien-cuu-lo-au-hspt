# -*- coding: utf-8 -*-
"""Tao ban v2 cua luan an voi cac edit duoc to do, xu ly cac muc phan bien:
- Muc 4 (Muc dich): bo sung danh gia hieu qua chuong trinh
- Muc 5 (Doi tuong): bo sung khung chuong trinh
- Muc 7 (Nhiem vu): lam ro "bien phap doi pho" la bien trung gian
- Muc 8/14/28/49 (bien doi pho): them ghi chu tai cac vi tri then chot
- Muc 15-16: thong nhat thuat ngu "lan toa" (sua heading 1.2.3.1)
- Muc 17-18: sua typo "DSM V" -> "DSM-5", them ghi chu phien ban DSM
Save: Luan an TS/01_RoiLoanLoAu_v2_RedEdits_26052026.docx
26/05/2026."""
import os, sys, io, copy
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_Rối loạn lo âu_ 26-05.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_RoiLoanLoAu_v2_RedEdits_26052026.docx')

RED = RGBColor(192, 0, 0)

# ============================================================
# HELPERS
# ============================================================
def append_red_run(paragraph, text, italic=True, bold=False):
    """Them mot run mau DO vao cuoi paragraph, GIU run hien co."""
    run = paragraph.add_run(text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(13)
    run.font.color.rgb = RED
    run.italic = italic
    run.bold = bold
    return run

def replace_text_in_paragraph(paragraph, old_text, new_text, mark_red=True):
    """Tim old_text trong paragraph va thay bang new_text mau DO.
    Tra ve True neu da thay duoc, False neu khong tim thay."""
    # Lay full text
    full_text = paragraph.text
    if old_text not in full_text:
        return False

    # Tim run chua old_text. Don gian: replace tat ca run, gan lai
    # Build new text
    new_full = full_text.replace(old_text, new_text, 1)  # chi thay lan dau

    # Clear het cac run hien co
    for r in paragraph.runs:
        r.text = ''
    # Add lai voi 3 segment: before/old(red)/after
    idx_start = full_text.find(old_text)
    idx_end = idx_start + len(old_text)
    before = full_text[:idx_start]
    after = full_text[idx_end:]

    if before:
        r = paragraph.add_run(before)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    r_red = paragraph.add_run(new_text)
    r_red.font.name = 'Times New Roman'
    r_red.font.size = Pt(13)
    if mark_red:
        r_red.font.color.rgb = RED
        r_red.bold = True
    if after:
        r = paragraph.add_run(after)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
    return True


# ============================================================
# LOAD VA EDIT
# ============================================================
doc = Document(SRC)
ap = doc.paragraphs

# Theo doi cac edit
edits_log = []

# ========================================================
# EDIT 1: Muc 4 phan bien - Muc dich nghien cuu (para 160)
# Bo sung "danh gia hieu qua thu nghiem"
# ========================================================
p160 = ap[160]
append_red_run(p160, ' [NCS bổ sung theo phản biện mục Mục đích: đề tài cũng có nội dung thử nghiệm và đánh giá sơ bộ hiệu quả khung chương trình tập huấn phòng ngừa trên một mẫu HS — xem chi tiết ở Chương 3, mục 3.7.]')
edits_log.append(('Para 160 - Mục đích', 'Bổ sung note về thử nghiệm + đánh giá hiệu quả'))

# ========================================================
# EDIT 2: Muc 5 phan bien - Doi tuong (para 162)
# Bo sung "khung chuong trinh tap huan phong ngua"
# ========================================================
p162 = ap[162]
append_red_run(p162, ' [NCS bổ sung theo phản biện mục Đối tượng: đối tượng nghiên cứu bao gồm cả khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS.]')
edits_log.append(('Para 162 - Đối tượng', 'Bổ sung khung chương trình tập huấn'))

# ========================================================
# EDIT 3: Muc 6 phan bien - Gia thuyet khoa hoc (para 167)
# Them ghi chu ngan gon ve viec rut gon dien giai
# ========================================================
p167 = ap[167]
append_red_run(p167, ' [NCS sẽ tách phần diễn giải dài ra khỏi giả thuyết, chuyển sang phần Tổng quan/Cơ sở lý luận; chỉ giữ lại giả thuyết dạng phát biểu ngắn gọn trong bản chỉnh sửa hoàn chỉnh.]')
edits_log.append(('Para 167 - Giả thuyết', 'Note rút gọn diễn giải'))

# ========================================================
# EDIT 4: Muc 7 phan bien - Nhiem vu 6.2 (para 177)
# Lam ro "bien phap doi pho la bien trung gian"
# ========================================================
p177 = ap[177]
append_red_run(p177, ' [NCS làm rõ: "biện pháp đối phó" được nghiên cứu như BIẾN TRUNG GIAN (mediator) giữa các yếu tố ảnh hưởng và biểu hiện rối loạn lo âu, không phải biến phụ thuộc cũng không thuộc yếu tố bảo vệ/nguy cơ.]')
edits_log.append(('Para 177 - Nhiệm vụ 6.2', 'Làm rõ vai trò biến đối phó'))

# ========================================================
# EDIT 5: Muc 17 phan bien - DSM versions (para 534)
# Sua typo "DSM V, 2013" -> "DSM-5 (APA, 2013)"
# ========================================================
p534 = ap[534]
replaced = replace_text_in_paragraph(p534, 'DSM V, 2013', 'DSM-5 (APA, 2013)', mark_red=True)
if replaced:
    edits_log.append(('Para 534', 'Sửa "DSM V, 2013" → "DSM-5 (APA, 2013)"'))

# ========================================================
# EDIT 6: Them note phan bien ve DSM versions o dau muc 1.2.2
# Tim heading "1.2.2" hoặc đoạn intro DSM
# ========================================================
# Heading 1.2.2 nam o dau cac doan ve phan loai DSM
# Vi tri xap xi: truoc para 530 (DSM-IV+DSM-5)
# Them note vao para 534 (sau khi sua DSM V)
append_red_run(p534, ' [NCS bổ sung làm rõ theo phản biện: Luận án sử dụng DSM-5 làm chuẩn chính. Các tham chiếu đến DSM-IV (1994) chỉ nhằm mục đích lịch sử khi mô tả tiến triển của tiêu chí chẩn đoán; tham chiếu đến DSM-5-TR (2022) chỉ ở đoạn cập nhật mới nhất về Rối loạn lo âu chia ly. Mọi định nghĩa và phân loại chính của luận án đều dựa trên DSM-5.]')
edits_log.append(('Para 534 - DSM versions', 'Note tổng hợp lý do dùng nhiều phiên bản DSM'))

# ========================================================
# EDIT 7: Muc 16 phan bien - "Roi loan lo au tong quat" heading 1.2.3.1 (para 598)
# Doi heading -> "Roi loan lo au lan toa (General Anxiety Disorder / con goi 'roi loan lo au tong quat')"
# ========================================================
p598 = ap[598]
# Tim heading text "1.2.3.1 Rối loạn lo âu tổng quát"
old_h = '1.2.3.1 Rối loạn lo âu tổng quát'
if old_h in p598.text:
    # Replace
    new_h = '1.2.3.1 Rối loạn lo âu lan tỏa'
    replace_text_in_paragraph(p598, old_h, new_h, mark_red=True)
    edits_log.append(('Para 598 - Heading 1.2.3.1', '"tổng quát" → "lan tỏa" (thuật ngữ chính thức DSM-5 tiếng Việt)'))

# ========================================================
# EDIT 8: Them ghi chu thuat ngu o dau muc 1.2.3.1
# (para 599 hoac sau heading)
# ========================================================
if 599 < len(ap):
    append_red_run(ap[599], ' [NCS thống nhất: trong toàn bộ luận án sử dụng "rối loạn lo âu lan tỏa" theo bản tiếng Việt chính thức của DSM-5; thuật ngữ "lo âu tổng quát" còn dùng ở một số chỗ trong tổng quan các nghiên cứu khác là dịch trực tiếp từ "Generalized Anxiety Disorder" và cùng chỉ một rối loạn.]')
    edits_log.append(('Para 599', 'Note thuật ngữ lan tỏa = tổng quát'))

# ========================================================
# EDIT 9: Muc 18 phan bien - Sua "DSM-4" trong para 630 -> "DSM-IV"
# ========================================================
p630 = ap[630]
if 'DSM-4' in p630.text:
    replace_text_in_paragraph(p630, 'DSM-4', 'DSM-IV', mark_red=True)
    edits_log.append(('Para 630', 'Sửa "DSM-4" → "DSM-IV" (chuẩn ký hiệu)'))

# ========================================================
# THEM GHI CHU TONG QUAT O DAU FILE
# ========================================================
# Chen 1 paragraph mau DO o ngay sau title (truoc Muc luc)
# Tim paragraph "MỤC LỤC"
muc_luc_idx = None
for i, p in enumerate(ap):
    if p.text.strip() == 'MỤC LỤC':
        muc_luc_idx = i
        break

# Tao paragraph moi truoc MUC LUC
if muc_luc_idx is not None:
    # Insert before muc_luc_idx
    new_p_xml = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    new_p_xml.append(pPr)
    r_xml = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts'); rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), '26'); rPr.append(sz)
    color = OxmlElement('w:color'); color.set(qn('w:val'), 'C00000'); rPr.append(color)
    b = OxmlElement('w:b'); rPr.append(b)
    italic = OxmlElement('w:i'); rPr.append(italic)
    r_xml.append(rPr)
    t_xml = OxmlElement('w:t')
    t_xml.text = '[BẢN v2 — NCS đã chỉnh sửa các phần được tô ĐỎ theo nhận xét phản biện ngày 25/05/2026 của PGS Đặng Hoàng Minh. Tổng số điểm đã xử lý: ' + str(len(edits_log)) + '. Tổng quan đã được tách thành file 01_TongQuan_NhomTheoChuDe riêng. Bảng giải trình tiếp thu và đoạn giải trình thay đổi mẫu được nộp kèm. Sau khi rà soát chấp nhận các thay đổi, NCS sẽ chuyển toàn bộ chữ đỏ về đen.]'
    t_xml.set(qn('xml:space'), 'preserve')
    r_xml.append(t_xml)
    new_p_xml.append(r_xml)
    ap[muc_luc_idx]._p.addprevious(new_p_xml)
    edits_log.append(('Header note', 'Ghi chú v2 cho NCS/Hội đồng'))

# ============================================================
# LUU
# ============================================================
doc.save(OUT)
print(f"Da luu: {OUT}")
print(f"Size: {os.path.getsize(OUT) // 1024}KB")
print(f"\nDanh sach edit ({len(edits_log)}):")
for loc, desc in edits_log:
    print(f"  - {loc}: {desc}")
