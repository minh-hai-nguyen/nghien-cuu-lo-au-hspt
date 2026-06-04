# -*- coding: utf-8 -*-
"""Tao 6 tom tat cho QT012-QT017 theo format CTH v5."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name='Times New Roman'; s.font.size=Pt(12)
    s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.5
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3); sec.right_margin=Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
def save(doc, name):
    p = os.path.join(TT_DIR, name); doc.save(p)
    d = Document(p); chars = sum(len(x.text) for x in d.paragraphs)
    print(f'  {name}: {chars} chars, {len(d.tables)} tables')

# ============================================================
# QT012 — GBD ASEAN 2025 Lancet Public Health
# ============================================================
print('[1/6] QT012 GBD ASEAN')
doc = make_doc()
H(doc, 'Tóm tắt bài QT012 — Gánh nặng 10 rối loạn tâm thần tại ASEAN 1990–2021', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Epidemiology and burden of 10 mental disorders across the Association of '
       'Southeast Asian Nations, 1990–2021: findings from the Global Burden of Disease Study 2021" '
       'của GBD 2021 ASEAN Collaborators (2025), đăng trên The Lancet Public Health vol. 10, '
       'tr. e480, 06/2025. DOI 10.1016/S2468-2667(25)00098-2. Đây là phân tích đại diện cho '
       '10 quốc gia ASEAN (Brunei, Campuchia, Indonesia, Lào, Malaysia, Myanmar, Philippines, '
       'Singapore, Thái Lan, Việt Nam) về dịch tễ 10 rối loạn tâm thần, sử dụng dữ liệu GBD 2021 '
       'với DisMod-MR 2.1 (hồi quy siêu phân tích Bayes).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Phân tích dữ liệu GBD 2021 cho 10 rối loạn: trầm cảm, lo âu, lưỡng cực, tâm thần phân '
       'liệt, tự kỷ, rối loạn hành vi, ADHD, rối loạn ăn uống, khuyết tật trí tuệ phát triển tự '
       'phát và các rối loạn tâm thần khác. Định nghĩa ca bệnh theo DSM hoặc ICD. Các chỉ số '
       'chính: YLDs, YLLs, DALYs với khoảng bất định 95% UIs. Phân tích theo tuổi, giới tính, '
       'năm (1990–2021) và địa điểm (10 quốc gia ASEAN).')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Vào năm 2021, tổng cộng có 80,4 triệu (UI 73,8–87,2) trường hợp rối loạn tâm thần '
       'tại ASEAN — tăng 70% kể từ 1990. Gánh nặng DALYs đạt 11,2 triệu (UI 8,54–14,3) năm '
       'sống điều chỉnh theo khuyết tật. Tỷ lệ hiện mắc rối loạn tâm thần khác nhau giữa các '
       'quốc gia: Việt Nam 10,1%, Malaysia 13,2%.', italic=True)
P(doc, 'Phát hiện đáng chú ý: Trong nhóm tuổi 10–14, rối loạn tâm thần chiếm 16,3% tổng DALYs '
       '— ghi nhận VTN châu Á là nhóm dễ tổn thương đặc biệt.', bold=True)
table(doc, ['Chỉ số', 'Giá trị 2021', 'Ghi chú'],
      [['Tổng số ca mắc', '80,4 triệu (73,8–87,2)', '↑ 70% so với 1990'],
       ['Tổng DALYs', '11,2 triệu (8,54–14,3)', 'Chiếm gánh nặng lớn ở VTN'],
       ['Tỷ lệ Việt Nam', '10,1 %', 'Thấp hơn Malaysia (13,2 %)'],
       ['Nhóm 10–14 tuổi', '16,3 % DALYs', 'VTN là nhóm dễ tổn thương']],
      widths=[5.0, 5.0, 5.5])

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Đây là phân tích GBD 2021 CẬP NHẬT NHẤT cho khu vực ASEAN — 10 quốc gia, '
       'giai đoạn 1990–2021, sử dụng DisMod-MR 2.1 là mô hình hoá tiên tiến nhất. Lancet Public '
       'Health Q1 IF ≈ 40. Cung cấp bằng chứng cơ sở CHO CHÍNH SÁCH SKTT khu vực. Tác động '
       'chính sách: xác định VTN 10–14 tuổi là nhóm ưu tiên.')
P(doc, 'Hạn chế: (1) Dữ liệu đầu vào cho ASEAN (đặc biệt VN, Lào, Campuchia, Myanmar, Brunei) '
       'còn hạn chế, dẫn đến uncertainty intervals rộng; (2) Không tách biệt rối loạn lo âu '
       'với trầm cảm theo chi tiết; (3) Chưa phân tích các yếu tố quyết định (determinants); '
       '(4) Phụ thuộc vào chất lượng dữ liệu ICD/DSM địa phương — các nước như Lào, Myanmar '
       'có hệ thống báo cáo yếu.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần bổ sung dữ liệu sơ cấp (khảo sát đại diện quốc gia) cho VN, Lào, Campuchia — hiện '
       'VN có V-NAMHS 2022 là nguồn tốt nhất. Tách biệt phân tích lo âu vs trầm cảm. Phân tích '
       'các yếu tố quyết định xã hội (SDH) và đánh giá hiệu quả chính sách SKTT khu vực.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao. Lancet Q1 IF ≈ 40, phương pháp GBD chuẩn, '
       'dữ liệu khu vực cập nhật nhất, ý nghĩa chính sách cao.', bold=True)
save(doc, 'QT012_GBD_ASEAN_2025_Lancet.docx')

# ============================================================
# QT013 — Zhameden (Yin) 2025 PLOS ONE LMIC School Interventions
# ============================================================
print('[2/6] QT013 Zhameden LMIC')
doc = make_doc()
H(doc, 'Tóm tắt bài QT013 — Can thiệp trường học phòng ngừa lo âu và trầm cảm tại LMIC', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "School-based interventions to prevent anxiety and depression in children '
       'and adolescents in low- and middle-income countries: A systematic review" của Yin SZD, '
       'Low MK, Mishu MP (2025, University of York UK), đăng trên PLOS ONE vol. 20(4), e0316825. '
       'DOI 10.1371/journal.pone.0316825. Open Access. Tổng quan hệ thống các RCT can thiệp '
       'phòng ngừa lo âu và trầm cảm tại trường học ở các nước thu nhập thấp–trung bình (LMIC).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống Ovid MEDLINE, Embase, PsycINFO, CENTRAL cho bài báo xuất bản từ '
       '2018 đến tháng 7/2023. Bao gồm các RCT đánh giá can thiệp trường học cho trẻ em/VTN '
       '4–18 tuổi tại LMIC. Chỉ nhận nghiên cứu tiếng Anh. Kết quả chính: triệu chứng lo âu '
       'và/hoặc trầm cảm. Đánh giá nguy cơ thiên vị (RoB) cho tất cả RCT được đưa vào.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Từ 3.863 bài xác định được, chỉ 6 nghiên cứu với 1.587 học sinh đáp ứng tiêu chí. '
       'Trong 4 nghiên cứu đánh giá phòng ngừa CẢ lo âu và trầm cảm: CHỈ 1 NGHIÊN CỨU cho thấy '
       'giảm triệu chứng LO ÂU có ý nghĩa. Trong 4 nghiên cứu đánh giá trầm cảm: 3/4 NGHIÊN CỨU '
       'báo cáo cải thiện trầm cảm có ý nghĩa. Tất cả nghiên cứu đều được phân loại có NGUY CƠ '
       'THIÊN VỊ CAO hoặc có một số lo ngại (GRADE very low).', bold=True)

table(doc, ['Outcome', 'Số NC', 'Cải thiện', 'Kết luận'],
      [['Lo âu', '4', '1/4 (25%)', 'Không rõ hiệu quả'],
       ['Trầm cảm', '4', '3/4 (75%)', 'Có tiềm năng'],
       ['Chất lượng bằng chứng', '6', 'Tất cả RoB cao', 'GRADE rất thấp']],
      widths=[5.0, 2.5, 3.5, 4.5])

P(doc, 'LƯU Ý QUAN TRỌNG: 0 nghiên cứu từ Việt Nam trong 6 RCT được tổng hợp — phù hợp với '
       'khẳng định "0 RCT VN cho VTN" trong v2, nhưng cần được cập nhật với Happy House '
       'Tran 2023 là RCT đầu tiên của VN.', italic=True)

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: SR đầu tiên tập trung vào LMIC cho phòng ngừa lo âu/trầm cảm tại trường — '
       'rất phù hợp bối cảnh Việt Nam. PLOS ONE Open Access. Đánh giá RoB nghiêm ngặt. Thông '
       'báo rõ: chất lượng bằng chứng hiện tại THẤP.')
P(doc, 'Hạn chế: (1) Chỉ 6 RCT — quá ít để rút kết luận chắc chắn; (2) Chỉ 2018–2023 — có thể '
       'bỏ lỡ nghiên cứu cũ; (3) Chỉ tiếng Anh — bỏ lỡ bài tiếng Trung, Tây Ban Nha; (4) Không '
       'tính Happy House VN 2023 (có thể đã nằm trong pipeline khi bài được viết); (5) Không '
       'phân tích theo đặc điểm can thiệp (universal vs targeted, duration).')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần RCT chất lượng cao hơn với RoB thấp tại LMIC. Việt Nam cần RCT chuyên biệt cho '
       'LO ÂU (không chỉ trầm cảm — Happy House chủ yếu đo trầm cảm). So sánh universal vs '
       'targeted. Tích hợp với chương trình trường học sẵn có.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SR đầu tiên LMIC, Q1 Open Access, quan trọng cho '
       'định hướng nghiên cứu VN. Hạn chế: bằng chứng nền mỏng.', bold=True)
save(doc, 'QT013_Zhameden_2025_PLOSONE_LMIC.docx')

# ============================================================
# QT014 — Anderson 2025 Wiley Narrative
# ============================================================
print('[3/6] QT014 Anderson Narrative')
doc = make_doc()
H(doc, 'Tóm tắt bài QT014 — Các yếu tố góp phần gia tăng lo âu ở VTN: Tổng quan tường thuật', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Contributing Factors to the Rise in Adolescent Anxiety and Associated '
       'Mental Health Disorders: A Narrative Review of Current Literature" của Anderson, T.L., '
       'Valiauga, R., Tallo, C., Hong, C.B., Manoranjithan, S., Domingo, C., Pilitsis, J.G. '
       'và cộng sự (2025), đăng trên Journal of Child and Adolescent Psychiatric Nursing vol. 38, '
       'e70009. DOI 10.1111/jcap.70009. Tác giả từ Mỹ (Rowan University School of Osteopathic '
       'Medicine). Là tổng quan tường thuật 61 bài về yếu tố góp phần gia tăng lo âu VTN.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tổng quan tường thuật (narrative review) — KHÔNG phải SR/MA. Tìm kiếm PubMed, '
       'ScienceDirect, Medline. Chọn các bài gốc và review thảo luận về gia tăng lo âu ở '
       'Thế hệ Z (sinh 1997–2012). Phân tích theo 4 nhóm yếu tố: áp lực học tập, ảnh hưởng MXH, '
       'động lực gia đình, các yếu tố gây căng thẳng xã hội rộng lớn. Đề xuất can thiệp 3 cấp: '
       'lâm sàng, giáo dục, cộng đồng.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Phát hiện chính:')
P(doc, '• Tỷ lệ lo âu VTN 13–18 tuổi là 31,9 % (Daly 2022 — trích dẫn) — Thế hệ Z CAO NHẤT '
       'so với 3 thế hệ trước đó.')
P(doc, '• 48/52 nghiên cứu xác nhận tương quan áp lực học tập → SKTT kém (Steare và cộng sự '
       '2023 — trích dẫn).')
P(doc, '• MXH ảnh hưởng lo âu qua nhiều cơ chế: so sánh xã hội, bắt nạt mạng, FOMO, thiếu ngủ.')
P(doc, '• Gia đình: cha mẹ lo âu truyền sang con (trans-generational transmission), kỳ vọng '
       'quá cao, ít tương tác cá nhân.')
P(doc, '• Yếu tố xã hội: COVID-19, bất ổn kinh tế, biến đổi khí hậu, thông tin về khủng hoảng '
       'toàn cầu.')

P(doc, 'Đề xuất can thiệp đa tầng:', bold=True)
table(doc, ['Tầng can thiệp', 'Biện pháp', 'Ví dụ'],
      [['Lâm sàng', 'CBT, SSRI, therapist-guided', 'CAMS Walkup 2008'],
       ['Giáo dục', 'Trường học, kỹ năng ứng phó', 'Happy House VN 2023'],
       ['Cộng đồng', 'Chính sách SKTT, giảm kỳ thị', 'WHO 2022 framework']],
      widths=[3.5, 6.0, 5.5])

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Tổng quan 61 bài bao quát nhiều yếu tố, viết dễ hiểu, phù hợp đào tạo. '
       'Tiếp cận đa tầng rõ ràng. Journal of Child and Adolescent Psychiatric Nursing là tạp '
       'chí ứng dụng tốt.')
P(doc, 'Hạn chế: (1) NARRATIVE REVIEW, KHÔNG phải SR/MA — không có tìm kiếm hệ thống hay đánh '
       'giá chất lượng chặt chẽ (Green và cộng sự 2006); (2) Dữ liệu chủ yếu từ phương Tây '
       '(Mỹ, UK, Úc) — có thể không generalizable cho châu Á (Polanczyk 2015); (3) Không có '
       'phân tích định lượng — chỉ mô tả; (4) Khả năng thiên vị tác giả cao do tự chọn bài.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần SR + MA (không phải narrative) để đánh giá định lượng hiệu quả các yếu tố. Nghiên '
       'cứu dọc theo Thế hệ Z VN. So sánh Việt Nam với các nước phương Tây về các yếu tố tương '
       'đồng và khác biệt.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐ Trung bình–Khá. Tổng quan tường thuật — hữu ích để có bức '
       'tranh tổng thể nhưng thiếu bằng chứng định lượng.', bold=True)
save(doc, 'QT014_Anderson_2025_Wiley_Narrative.docx')

# ============================================================
# QT015 — Zhu 2025 BMC China
# ============================================================
print('[4/6] QT015 Zhu 2025')
doc = make_doc()
H(doc, 'Tóm tắt bài QT015 — Yếu tố ảnh hưởng trầm cảm HS THCS/THPT miền Đông Trung Quốc', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Influencing factors and changing trends of depressive symptoms in middle '
       'and high school students in East China from 2019 to 2023" của Zhu và cộng sự (2025), '
       'đăng trên BMC Public Health vol. 25, article 17. DOI 10.1186/s12889-024-21252-8. '
       'Khảo sát cắt ngang trên 9.831 học sinh THCS và THPT tại Suzhou, miền Đông Trung Quốc. '
       'Sử dụng PHQ-9 (Patient Health Questionnaire-9 mục) đánh giá trầm cảm.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Khảo sát cắt ngang lặp lại 2019, 2021, 2023 (trước, trong và sau COVID) trên 9.831 '
       'HS tại Suzhou. Sử dụng PHQ-9 với 2 ngưỡng cắt: ≥5 "có thể có trầm cảm" (possible), ≥10 '
       '"chắc chắn có trầm cảm" (definite). Phân tích hồi quy logistic đa biến xác định yếu tố '
       'nguy cơ và bảo vệ.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ: 14,5 % HS có triệu chứng trầm cảm "có thể" và 5,8 % "chắc chắn". Tổng 20,3 % '
       'có triệu chứng trầm cảm — tương đương với nhiều NC khác ở Trung Quốc.')

table(doc, ['Yếu tố', 'AOR', 'KTC 95 %', 'Diễn giải'],
      [['Ngủ < 5 giờ (chắc chắn)', '13,710', '—', 'NGUY CƠ RẤT MẠNH NHẤT'],
       ['Gia đình tái hôn', '1,837', '—', 'Nguy cơ cao'],
       ['Gia đình đơn thân', '1,434', '—', 'Nguy cơ'],
       ['HS THPT (vs THCS)', '1,409', '—', 'Nguy cơ'],
       ['Nam giới (vs nữ)', '0,803', '—', 'Bảo vệ'],
       ['Hoạt động ngoài trời > 1h/ngày', '0,666–0,785', '—', 'Bảo vệ rõ ràng']],
      widths=[5.0, 2.0, 2.5, 6.0])

P(doc, 'PHÁT HIỆN THEN CHỐT: Ngủ < 5 giờ có AOR = 13,71 cho "trầm cảm chắc chắn" — yếu tố '
       'nguy cơ mạnh nhất được phát hiện. Đây là con số có ý nghĩa thực hành cao cho can '
       'thiệp — cải thiện giấc ngủ có thể là chiến lược phòng ngừa hiệu quả nhất.', bold=True)

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Cỡ mẫu rất lớn (n = 9.831), so sánh 3 thời điểm 2019–2023 bao gồm cả giai '
       'đoạn COVID. BMC Public Health Q1. Phát hiện AOR=13,71 cho giấc ngủ < 5h là CON SỐ '
       'ĐÁNG NHỚ — dễ ứng dụng cho can thiệp (cải thiện giấc ngủ + hoạt động ngoài trời).')
P(doc, 'Hạn chế: (1) Chỉ tại Suzhou (miền Đông, đô thị phát triển) — không đại diện nông thôn '
       'hay miền Tây TQ; (2) PHQ-9 là sàng lọc, không phải chẩn đoán lâm sàng; (3) Chỉ đo trầm '
       'cảm, thiếu lo âu (cần GAD-7 song song); (4) 3 thời điểm là 3 nhóm khác nhau, không '
       'phải cùng nhóm theo dõi dọc.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Lặp lại tại VN với HS THCS/THPT — đặc biệt kiểm tra vai trò giấc ngủ (VN có Hoàng '
       'Trung Học 2025 đã ghi nhận giấc ngủ β = −0,149). Nghiên cứu dọc theo dõi cùng nhóm '
       'qua COVID. RCT can thiệp tăng giấc ngủ + hoạt động ngoài trời tại trường.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. BMC Q1, mẫu lớn, phát hiện AOR=13,71 giấc ngủ rất '
       'quan trọng cho thực hành.', bold=True)
save(doc, 'QT015_Zhu_2025_BMC_China.docx')

# ============================================================
# QT016 — Mudunna 2025 Lancet SEA South Asia
# ============================================================
print('[5/6] QT016 Mudunna South Asia')
doc = make_doc()
H(doc, 'Tóm tắt bài QT016 — SKTT VTN Nam Á: Tổng quan hệ thống', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Nature, prevalence, and determinants of mental health problems among South '
       'Asian adolescents: A systematic review" của Mudunna và cộng sự (2025), đăng trên The '
       'Lancet Regional Health — Southeast Asia 2025. DOI 10.1016/S2772-3682(25)00003-4. '
       'Tổng quan hệ thống về bản chất, tỷ lệ và các yếu tố quyết định của các vấn đề SKTT ở '
       'vị thành niên (10–19 tuổi) tại 8 nước Nam Á. Nam Á chiếm 24 % dân số thế giới, chủ yếu '
       'là các nước thu nhập thấp và trung bình (LMIC).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống trên MEDLINE, PsycINFO, Embase, CINAHL, Global Health theo hướng dẫn '
       'PRISMA. Tiêu chí: nghiên cứu về VTN 10–19 tuổi tại 8 nước Nam Á (Ấn Độ, Pakistan, '
       'Bangladesh, Nepal, Sri Lanka, Bhutan, Maldives, Afghanistan). Đánh giá chất lượng '
       'phương pháp. KHÔNG thực hiện meta-analysis do tính không đồng nhất quá cao.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ hiện mắc vấn đề SKTT dao động RẤT LỚN từ 1,5 % đến 81,6 % — do khác biệt về '
       'phương pháp và quốc gia. Trầm cảm và lo âu là phổ biến nhất.')

P(doc, 'Các yếu tố quyết định chính:', bold=True)
table(doc, ['Nhóm yếu tố', 'Các yếu tố cụ thể'],
      [['Giới tính', 'Nữ giới có nguy cơ cao hơn nam'],
       ['Kinh tế – xã hội', 'Nghèo đói, bất bình đẳng'],
       ['Gia đình', 'Bạo lực gia đình, mất cha mẹ, thiếu hỗ trợ xã hội'],
       ['Trường học', 'Bắt nạt, áp lực học tập'],
       ['Văn hoá', 'Kỳ thị SKTT, rào cản văn hoá ngăn tiếp cận dịch vụ']],
      widths=[4.0, 11.5])

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Tổng quan đầu tiên tập trung vào VTN Nam Á — khu vực chiếm 24 % dân số '
       'thế giới nhưng ít được nghiên cứu. Đăng trên Lancet Regional Health — tạp chí uy tín '
       'cao. Bao phủ 8 quốc gia với đa dạng văn hoá, tôn giáo, kinh tế. Xác định rõ khoảng '
       'trống: tỷ lệ dao động 1,5–81,6 % do thiếu chuẩn hoá phương pháp.')
P(doc, 'Hạn chế: (1) KHÔNG có meta-analysis do heterogeneity quá cao; (2) Phần lớn nghiên cứu '
       'từ Ấn Độ — thiếu dữ liệu từ Afghanistan, Bhutan, Maldives; (3) Tập trung Nam Á — không '
       'bao gồm Đông Nam Á (VN, Indonesia, Philippines); (4) Không phân biệt rõ lo âu và trầm '
       'cảm trong báo cáo tổng hợp.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Mở rộng sang Đông Nam Á để so sánh liên khu vực (phối hợp với GBD 2021 ASEAN '
       'Collaborators 2025 = QT012). Chuẩn hoá công cụ đo lường xuyên quốc gia Nam Á. RCT '
       'can thiệp phù hợp văn hoá tại trường học Nam Á. Việt Nam có thể học được từ mô hình '
       'De Silva Sri Lanka (QT038) về CBT trường học LMIC Nam Á.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SR đầu tiên Nam Á, Lancet Regional Health Q1, phản '
       'ánh khoảng trống quan trọng. Hạn chế: heterogeneity quá cao.', bold=True)
save(doc, 'QT016_Mudunna_2025_LancetSEA_SouthAsia.docx')

# ============================================================
# QT017 — Puyat 2025 Filipino Youth
# ============================================================
print('[6/6] QT017 Puyat Filipino')
doc = make_doc()
H(doc, 'Tóm tắt bài QT017 — Gia tăng trầm cảm thanh niên Philippines 2013 – 2021', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Rising depressive symptoms and widening sociodemographic inequalities '
       'among Filipino young people: findings from two nationally representative cross-sectional '
       'surveys" của Puyat và cộng sự (2025), đăng trên Cambridge Prisms: Global Mental Health '
       '2025. DOI 10.1017/gmh.2025.39. Phân tích dữ liệu đại diện quốc gia từ 2 khảo sát Thanh '
       'niên và Vị thành niên Philippines (YAFS — Young Adult Fertility and Sexuality Survey): '
       'YAFS 2013 (n = 19.178) và YAFS 2021 (n = 10.949), tổng cộng n = 30.127.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Phân tích dữ liệu thứ cấp từ 2 khảo sát đại diện quốc gia Philippines. Sử dụng '
       'CES-D-11 (Center for Epidemiologic Studies Depression Scale, phiên bản 11 mục) để ước '
       'tính tỷ lệ triệu chứng trầm cảm trung bình đến nặng (MSDS — Moderate to Severe '
       'Depressive Symptoms). So sánh 2013 vs 2021 theo các nhóm nhân khẩu xã hội: giới tính, '
       'tầng kinh tế, trình độ học vấn, SOGIE (Sexual Orientation, Gender Identity and '
       'Expression).')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Phát hiện KHÔNG MẤY LẠC QUAN: Tỷ lệ MSDS TĂNG HƠN GẤP ĐÔI từ 9,6 % (2013) lên 20,9 % '
       '(2021) — tăng 118 % trong 8 năm. Bất bình đẳng nhân khẩu xã hội MỞ RỘNG đáng kể.', bold=True)

table(doc, ['Nhóm dân số', '2013 (%)', '2021 (%)', 'Mức tăng'],
      [['Tổng thể', '9,6 %', '20,9 %', '× 2,2'],
       ['Nữ', '10,8 %', '24,3 %', '× 2,3'],
       ['Không thuộc giới chuẩn tắc (LGBTQ+)', '9,7 %', '32,3 %', '× 3,3'],
       ['Nhóm nghèo nhất', '12,3 %', '25,1 %', '× 2,0'],
       ['Không đi học', '8,5 %', '26,5 %', '× 3,1']],
      widths=[6.0, 3.0, 3.0, 3.5])

P(doc, 'Các nhóm dễ tổn thương nhất (LGBTQ+, không đi học) có mức tăng cao nhất — bất bình '
       'đẳng mở rộng trong gần 1 thập kỷ.', italic=True)

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Cỡ mẫu rất lớn (n = 30.127 tổng cộng), đại diện quốc gia Philippines — '
       'hiếm có ở Đông Nam Á. So sánh 2 thời điểm 2013 vs 2021 cho thấy xu hướng tăng gấp đôi '
       '— bằng chứng mạnh. Phân tích theo SOGIE rất tiến bộ, ít nghiên cứu ĐNA nào làm. Sử '
       'dụng khung SDH (Social Determinants of Health).')
P(doc, 'Hạn chế: (1) CES-D-11 sàng lọc, không chẩn đoán — tỷ lệ thực có thể khác nhiều; '
       '(2) Hai khảo sát khác nhau (YAFS4 2013 vs YAFS5 2021) — không phải theo dõi cùng '
       'nhóm, có thể có khác biệt phương pháp; (3) Chỉ đo trầm cảm, không đo lo âu; (4) '
       'Philippines có đặc thù văn hoá Kitô giáo + dân số trẻ — có thể không generalizable '
       'cho VN (Phật giáo + dân số già hơn).')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Bổ sung GAD-7 đo lo âu song song với CES-D trong khảo sát tiếp theo. Nghiên cứu dọc '
       'theo dõi cùng nhóm thanh niên qua nhiều năm. So sánh với VN (V-NAMHS 2022), Indonesia, '
       'Thái Lan để đánh giá xu hướng khu vực ASEAN. Lưu ý về SOGIE — VN cũng nên phân tích '
       'nhóm này để đánh giá bất bình đẳng.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Global Mental Health (Cambridge Q1), mẫu lớn, so '
       'sánh 2 thời điểm, phân tích SDH tiến bộ. Phát hiện LGBTQ+ × 3,3 đặc biệt đáng chú ý.',
  bold=True)
save(doc, 'QT017_Puyat_2025_Filipino_Youth.docx')

print('\n=== DONE 6 summaries ===')
