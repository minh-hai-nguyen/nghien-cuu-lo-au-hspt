# -*- coding: utf-8 -*-
"""Part 4: SERVICE USE + COVID-19 chapters (pages 25-35)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_imgs/'

doc = Document(OUT)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))
def H4(t): return P(t, bold=True, italic=True, size=11)

def PageMarker(page_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'--- Trang {page_num}, V-NAMHS, UNICEF/IOS 2022 ---')
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

def AddImg(filename, caption_vn, caption_en, width_cm=11.0):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try: p.add_run().add_picture(path, width=Cm(width_cm))
    except: return
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# =============================================================
# SERVICE USE chapter (pages 25-31)
# =============================================================
PageMarker('25')
H2('SỬ DỤNG DỊCH VỤ (SERVICE USE)')

H3('Tổng quan')
P('Các triệu chứng rối loạn tâm thần thường khởi phát trong giai đoạn vị thành niên (WHO 2014), mang đến cơ hội duy nhất cho can thiệp sớm và dịch vụ SKTT điều chỉnh riêng cho VTN (Colizzi, Lasalvia and Ruggeri 2020). Điều trị và hỗ trợ hiệu quả có thể giảm hoặc giải quyết triệu chứng và cải thiện chức năng, trong khi trải nghiệm tích cực sớm với dịch vụ chăm sóc SKTT có thể thúc đẩy hành vi tìm kiếm trợ giúp trong tương lai, giảm gánh nặng cho cá nhân và hệ thống y tế đến tuổi trưởng thành (Schnyder et al. 2019).')

P('Thu thập thông tin về mô hình sử dụng dịch vụ có thể giúp xác định khoảng trống trong cung cấp dịch vụ, thông tin xem dịch vụ SKTT hiện có có phù hợp và dễ tiếp cận hay không, và giải quyết các rào cản chăm sóc. Điều này đặc biệt quan trọng đối với VTN — nhóm đối mặt với thách thức đặc biệt khi tiếp cận chăm sóc do áp lực gia đình, ảnh hưởng bạn bè và khó khăn chi trả (Person, Hagquist và Michelson 2017, WHO 2014). Hiểu biết về các thách thức cụ thể với VTN cho phép hành động trực tiếp để giảm rào cản, cải thiện sử dụng dịch vụ và kết quả chăm sóc cho các em có vấn đề hay rối loạn SKTT.')

P('Tại VN, đã có những nỗ lực gần đây của Chính phủ nhằm tăng tiếp cận dịch vụ SKTT. Ví dụ, Thủ tướng Chính phủ ban hành Quyết định 1929/QĐ-TTg ngày 25/11/2020 (Thủ tướng Chính phủ 2020) phê duyệt Chương trình trợ giúp xã hội và phục hồi chức năng cho người tâm thần, trẻ em tự kỷ và người rối nhiễu tâm trí dựa vào cộng đồng giai đoạn 2021–2030. Tuy nhiên, VTN chỉ được nhắc đến như một trong nhiều nhóm đối tượng phụ, và tập trung vào người trẻ chủ yếu liên quan đến trẻ em tự kỷ. Gần đây hơn, Thủ tướng Chính phủ ban hành Quyết định 155/QĐ-TTg ngày 29/01/2022 (Thủ tướng Chính phủ 2022) phê duyệt "Kế hoạch quốc gia phòng chống bệnh không lây nhiễm và rối loạn sức khỏe tâm thần giai đoạn 2022–2025". Dù một trong các mục tiêu của Kế hoạch là quản lý điều trị ở cấp dân số để hạn chế tăng rối loạn tâm thần, trọng tâm vẫn là tâm thần phân liệt và trầm cảm, không chú ý tới chiến lược VTN-specific hay các rối loạn khác phổ biến trong giai đoạn VTN. Dù nhận thức về SKTT tại VN đang tăng, vẫn còn thiếu dịch vụ đáng kể, đặc biệt ở nông thôn so với đô thị, và tiếp cận còn bị tác động bởi tình trạng kinh tế – xã hội (Samuels et al. 2018).')

PageMarker('26')
P('Với VTN, các thách thức càng nghiêm trọng vì đa số dịch vụ SKTT hiện có chủ yếu phục vụ người lớn mắc bệnh tâm thần nặng và yêu cầu phức tạp như tâm thần phân liệt (Cao Tiến Đức 2020). Do đó, thiếu đáng kể dịch vụ SKTT cho VTN, đặc biệt cho những em trải qua suy giảm và căng thẳng từ triệu chứng rối loạn tâm thần nhưng không cần can thiệp chuyên sâu.')

P('Một trong các mục tiêu chính của V-NAMHS là xác định mức độ sử dụng dịch vụ SKTT ở VTN VN, cũng như mức nhu cầu được cảm nhận và rào cản chăm sóc. Chương này trình bày đo lường sử dụng dịch vụ và các yếu tố liên quan trong V-NAMHS và các phát hiện. Tất cả tỷ lệ và số lượng trong chương này đều được gia trọng.')

H3('Đo lường')
P('Câu hỏi về sử dụng dịch vụ được hỏi tất cả người tham gia, bất kể có ghi nhận triệu chứng rối loạn tâm thần hay không. Để hiểu liệu dịch vụ đã được sử dụng để hỗ trợ SKTT của VTN hay chưa, thuật ngữ "vấn đề cảm xúc và hành vi" (emotional and behavioural problems) được dùng để khung các câu hỏi dịch vụ. Thuật ngữ này được chọn để giảm kỳ thị hay tác động tiềm năng từ hiểu biết hạn chế về thuật ngữ SKTT. Ngoài ra, "vấn đề cảm xúc và hành vi" bao quát dải rộng các cách triệu chứng rối loạn tâm thần có thể biểu hiện, vốn khác nhau giữa các nhóm tuổi và bối cảnh văn hoá. Cách tiếp cận này nhất quán với các nghiên cứu tương tự về SKTT và sử dụng dịch vụ (Hafekost et al. 2016). Thêm nữa, V-NAMHS công nhận rằng dịch vụ được sử dụng có thể bao gồm nhà cung cấp vượt xa dịch vụ y tế chính thức. Vì vậy, một dải rộng nhà cung cấp dịch vụ được bao gồm trong các công cụ đo: y tế, giáo dục, tôn giáo/truyền thống, và các ngành khác. Tất cả câu hỏi liên quan đến 12 tháng qua và tất cả được hỏi phụ huynh, ngoại trừ hỗ trợ phi chính thức và chiến lược tự giúp (hỏi vị thành niên).')

H3('Phát hiện')
H4('Tần suất và loại sử dụng dịch vụ')
P('Trong số VTN có vấn đề SKTT (n = 1.301), chỉ 8,4 % đã sử dụng bất kỳ dịch vụ nào cung cấp hỗ trợ hoặc tư vấn cho vấn đề cảm xúc và hành vi trong 12 tháng qua. Tổng thể, chỉ 6,5 % (n = 389) VTN đã sử dụng dịch vụ trong 12 tháng qua, không có khác biệt giữa nam (7,4 %; n = 231) và nữ (5,5 %; n = 158). Trong số VTN đã sử dụng dịch vụ, hầu hết phụ huynh báo cáo dịch vụ có ích hoặc rất có ích (80,0 %).')

P('Về tần suất sử dụng dịch vụ trong 12 tháng qua, Bảng 12 cho thấy một nửa (50,8 %) chỉ sử dụng dịch vụ một lần, trong khi một phần tư (26,2 %) đã sử dụng 2 – 4 lần. Rất ít (3,4 %) VTN đã tiếp cận dịch vụ từ 5 lần trở lên.')

MakeTable(
    ['Giới', '% 1 lần', 'n', '% 2–4 lần', 'n', '% ≥5 lần', 'n'],
    [
        ('Nam', '47,2', '109', '26,3', '61', '4,5', '10'),
        ('Nữ', '56,0', '89', '26,2', '41', '1,7', '3'),
        ('Tổng', '50,8', '198', '26,2', '102', '3,4', '13'),
    ])
P('Bảng 12. Tần suất tiếp cận dịch vụ hỗ trợ/tư vấn vấn đề cảm xúc và hành vi trong 12 tháng qua ở VTN 10–17 tuổi theo giới. Gia trọng N: nam = 231; nữ = 158; tổng = 389. Nguồn: Bảng 12 báo cáo gốc, trang 26.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

PageMarker('27')
P('Khi được hỏi nhà cung cấp dịch vụ nào được dùng nhiều nhất trong 12 tháng qua, hơn một nửa phụ huynh (56,2 %) có con đã tiếp cận dịch vụ trong 12 tháng báo cáo đây là bác sĩ hoặc điều dưỡng (Bảng 13).')

MakeTable(
    ['Loại nhà cung cấp dịch vụ', '%', 'n'],
    [
        ('Bác sĩ hoặc điều dưỡng', '56,2', '219'),
        ('Nhân viên y tế cộng đồng', '10,7', '42'),
        ('Nhân viên trường học', '5,5', '22'),
        ('Lãnh đạo tôn giáo/tín ngưỡng', '4,5', '17'),
        ('Chuyên gia (bác sĩ tâm thần...)', '1,4', '5'),
        ('Khác', '0,3', '1'),
        ('Thầy thuốc truyền thống', '0,1', '0'),
    ])
P('Bảng 13. Loại nhà cung cấp dịch vụ được dùng nhiều nhất cho vấn đề cảm xúc và hành vi trong 12 tháng qua. Gia trọng N = 389. Nguồn: Bảng 13 báo cáo gốc, trang 27.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

AddImg('p027_img1_Image228.jpg',
       'Hình minh hoạ — Sử dụng dịch vụ SKTT',
       'Service use illustration',
       width_cm=10.0)

P('Trong số các em đã tiếp cận dịch vụ, 28,1 % có vấn đề SKTT trong 12 tháng qua, nghĩa là hơn 71,9 % không ghi nhận đủ triệu chứng để chỉ ra sự hiện diện của một vấn đề SKTT theo DISC-5. Như Hình 3 thể hiện, VTN có vấn đề SKTT đã tiếp cận dịch vụ có xu hướng báo cáo suy giảm nhiều hơn không, bất kể ghi nhận đầy đủ ngưỡng (14,8 %; n = 16) hay dưới ngưỡng (60,1 %; n = 66) triệu chứng.')

MakeTable(
    ['Loại triệu chứng', 'Tỷ lệ %'],
    [
        ('Triệu chứng đầy đủ + suy giảm (= rối loạn tâm thần)', '14,8'),
        ('Triệu chứng dưới ngưỡng + suy giảm', '60,1'),
        ('Triệu chứng đầy đủ không suy giảm', '0,4'),
        ('Triệu chứng dưới ngưỡng không suy giảm', '24,7'),
    ],
    header_bg='FFE6CC')
P('Hình 3. Ngưỡng triệu chứng và suy giảm ở VTN 10–17 tuổi có vấn đề SKTT đã sử dụng dịch vụ trong 12 tháng qua. Gia trọng N = 109. Nguồn: Figure 3 báo cáo gốc, trang 27.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H4('Nhu cầu được cảm nhận và rào cản chăm sóc')
P('Trong toàn bộ phụ huynh, chỉ 5,1 % (n = 307) xác định rằng con em họ cần trợ giúp cho vấn đề cảm xúc và hành vi. Trong số phụ huynh xác định con cần trợ giúp, 37,7 % (n = 116) báo cáo con đã nhận được tất cả trợ giúp cần thiết. Như Bảng 14 thể hiện, một phần năm phụ huynh (20,4 %) báo cáo ưa thích tự giải quyết vấn đề một mình hoặc với hỗ trợ gia đình; một phần mười báo cáo không chắc con có cần trợ giúp hay không (10,7 %) và/hoặc không biết nơi tìm trợ giúp (10,0 %). Lưu ý: phụ huynh có thể chọn nhiều lựa chọn trừ khi chọn rằng không có lý do nào trong danh sách áp dụng.')

PageMarker('28')
MakeTable(
    ['Lý do', '%', 'n'],
    [
        ('Ưa thích tự giải quyết vấn đề một mình hoặc với hỗ trợ gia đình', '20,4', '63'),
        ('Không chắc con có cần trợ giúp hay không', '10,7', '33'),
        ('Không biết nơi tìm trợ giúp', '10,0', '30'),
        ('Nghĩ vấn đề sẽ tự khỏi', '5,8', '18'),
        ('Không có nơi nào để tìm trợ giúp', '3,8', '12'),
        ('Chi phí quá cao hoặc gia đình không đủ khả năng', '2,7', '8'),
        ('Không có lý do nào trong danh sách (khác) ᵃ', '2,4', '7'),
        ('Khó khăn đi đến dịch vụ', '2,0', '6'),
        ('Lo ngại người khác sẽ nghĩ gì', '1,3', '4'),
        ('Không muốn nói chuyện với người lạ', '1,3', '4'),
        ('VTN từ chối trợ giúp hoặc không đến lịch hẹn', '0,7', '2'),
        ('Đã yêu cầu trợ giúp nhưng không nhận được', '0,6', '2'),
    ])
P('Bảng 14. Rào cản tìm kiếm hoặc nhận trợ giúp cho vấn đề cảm xúc và hành vi trong 12 tháng qua ở phụ huynh VTN 10–17 tuổi. Gia trọng N = 307; ᵃ chỉ tuỳ chọn duy nhất. Nguồn: Bảng 14 báo cáo gốc, trang 28.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H4('Hỗ trợ phi chính thức (Informal support)')
P('Tất cả VTN được hỏi thường nói chuyện với ai khi có lo lắng hoặc quan tâm. Các em có thể chọn nhiều lựa chọn trừ khi chọn không nói chuyện với ai (9,5 %). Trong số các em có nói chuyện với ai đó, gần ba phần tư (73,9 %) VTN nói chuyện với thành viên gia đình, tiếp theo là gần hai phần năm báo cáo nói chuyện với bạn bè (38,2 %) (Bảng 15).')

MakeTable(
    ['Người VTN thường nói chuyện khi có lo lắng', '%', 'n'],
    [
        ('Thành viên gia đình', '73,9', '4.429'),
        ('Bạn bè', '38,2', '2.288'),
        ('Giáo viên', '2,9', '176'),
        ('Bạn trai/bạn gái', '2,3', '137'),
        ('Thành viên cộng đồng', '1,6', '94'),
        ('Khác', '1,0', '60'),
        ('Bác sĩ', '0,5', '31'),
        ('Lãnh đạo tôn giáo/tín ngưỡng', '0,4', '22'),
    ])
P('Bảng 15. Người VTN 10–17 tuổi thường nói chuyện khi có lo lắng hoặc quan tâm. Nguồn: Bảng 15 báo cáo gốc, trang 28.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H4('Chiến lược tự giúp (Self-help strategies)')
P('Tất cả VTN được hỏi về chiến lược quản lý và phòng ngừa vấn đề cảm xúc và hành vi trong 12 tháng qua. VTN có thể chọn nhiều lựa chọn trừ khi báo cáo không sử dụng chiến lược tự giúp nào (11,7 %). Bảng 16 cho thấy 30,7 % ghi nhận tập thể dục nhiều hơn, trong khi một phần tư (24,8 %) nói họ làm nhiều hơn những điều họ thích.')

PageMarker('29')
MakeTable(
    ['Chiến lược tự giúp', '%', 'n'],
    [
        ('Tập thể dục nhiều hơn hoặc chơi thể thao', '30,7', '1.839'),
        ('Làm nhiều hơn những điều mình thích', '24,8', '1.490'),
        ('Tìm hỗ trợ từ gia đình', '19,2', '1.149'),
        ('Tìm hỗ trợ từ bạn bè', '17,7', '1.060'),
        ('Cải thiện chế độ ăn uống', '14,7', '884'),
        ('Tìm thông tin trong sách báo hoặc trên TV', '9,0', '538'),
        ('Tìm hỗ trợ qua mạng xã hội (chat room, social media, nhóm Internet...)', '6,9', '416'),
        ('Tham gia một nhóm xã hội nào đó', '1,7', '104'),
        ('Thiền hoặc tập thư giãn', '1,6', '97'),
        ('Ngừng hút thuốc, uống rượu hoặc dùng ma tuý', '0,2', '14'),
    ])
P('Bảng 16. Chiến lược tự giúp để quản lý hoặc phòng ngừa vấn đề cảm xúc và hành vi ở VTN 10–17 tuổi. Nguồn: Bảng 16 báo cáo gốc, trang 29.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H3('Suy xét (Considerations)')
H4('Diễn giải (Interpretation)')
P('Các phát hiện cho thấy nhu cầu không được đáp ứng có thể rất lớn đối với dịch vụ SKTT cho VTN tại VN. Trong nhóm VTN có vấn đề SKTT (21,7 %; n = 1.301), chỉ 8,4 % (n = 109) đã sử dụng bất kỳ dịch vụ hỗ trợ/tư vấn nào cho vấn đề cảm xúc và hành vi trong 12 tháng qua. Phát hiện này phù hợp với các khảo sát quốc tế chỉ ra rằng đa số người ở các nước thu nhập thấp và trung bình (LMICs) có rối loạn tâm thần không nhận được điều trị (Demyttenaere et al. 2004).')

P('Các phát hiện về rào cản chăm sóc có thể cung cấp một số manh mối lý giải vì sao tiếp cận dịch vụ thấp ở nhóm có vấn đề SKTT. Chỉ 5,1 % (n = 307) phụ huynh xác định con cần trợ giúp cho vấn đề cảm xúc và hành vi trong 12 tháng qua. Trong số đó, gần hai phần năm (37,7 %) báo cáo con đã nhận được tất cả trợ giúp cần thiết, và 20,4 % khác ưa thích tự giải quyết vấn đề với hỗ trợ gia đình. Dù một số VTN có thể đã nhận trợ giúp đủ, các phát hiện này cho thấy vấn đề về mental health literacy, đặc biệt trong nhận diện vấn đề SKTT cần trợ giúp chuyên nghiệp. Điều này được hỗ trợ bởi các nghiên cứu trước tại VN (Hoang-Minh Dang et al. 2020). Cũng có thể có ảnh hưởng của kỳ thị SKTT làm giảm khả năng tìm kiếm trợ giúp. Nghiên cứu trước đã tìm thấy kỳ thị SKTT ảnh hưởng hành vi tìm kiếm trợ giúp trong quần thể VN (Kamimura et al. 2018, Mai Do et al. 2014).')

P('Ngược lại, chỉ 28,1 % trong nhóm đã tiếp cận dịch vụ có vấn đề SKTT. Nghĩa là 71,9 % tiếp cận dịch vụ cho vấn đề cảm xúc và hành vi không ghi nhận đủ triệu chứng để chỉ ra vấn đề SKTT theo DISC-5. Phát hiện này khó diễn giải do tỷ lệ tiếp cận dịch vụ nhìn chung đã thấp (6,5 %; n = 389). Có thể một phần VTN đã nhận trợ giúp đủ từ dịch vụ để giảm triệu chứng đến mức không còn đáp ứng ngưỡng vấn đề SKTT. Tuy nhiên, do DISC-5 đo triệu chứng trong 12 tháng qua và một nửa VTN chỉ tiếp cận dịch vụ một lần, điều này dường như không xảy ra trong mọi trường hợp. Cũng có thể một số em có vấn đề SKTT nhưng không được nhận diện dựa trên triệu chứng đánh giá bởi DISC-5, hoặc khái niệm "vấn đề cảm xúc và hành vi" không tương quan tốt với vấn đề SKTT như dự định. Thêm nữa, dù mọi nỗ lực đã được thực hiện để thích ứng DISC-5 với bối cảnh VN, vẫn có thể thuật toán chấm điểm DSM-5 cần rà soát thêm để bắt chính xác các phân bố triệu chứng trong quần thể VN. Dữ liệu V-NAMHS cho phép nghiên cứu này trong tương lai.')

PageMarker('30')
H4('Hạn chế (Limitations)')
P('Các phát hiện V-NAMHS về sử dụng dịch vụ cho vấn đề cảm xúc và hành vi nên được diễn giải thận trọng. Thứ nhất, số lượng nhỏ VTN tiếp cận dịch vụ (n = 389) khiến khó xác định mô hình rõ ràng (ví dụ: tần suất hoặc sự hiện diện của vấn đề SKTT). Tuy nhiên, các phát hiện này là điểm khởi đầu và ở mức rộng nhất, chứng minh nhu cầu không được đáp ứng và rào cản chăm sóc trong quần thể VTN VN. Thứ hai, hầu hết câu hỏi được hỏi phụ huynh thay vì VTN do cân nhắc về độ dài phỏng vấn tổng thể. Điều này nghĩa là quan điểm của VTN, đặc biệt về rào cản chăm sóc, không được đo. Tuy nhiên, vì hầu hết VTN báo cáo nói chuyện với thành viên gia đình khi có lo lắng, vai trò phụ huynh như "người gác cổng" để tiếp cận dịch vụ SKTT cho thấy họ có vị trí tốt để cung cấp thông tin về rào cản liên quan nhất với hoạch định chính sách. Thêm nữa, câu hỏi dịch vụ hỏi phụ huynh có thể dễ điều chỉnh cho người trả lời VTN trong nghiên cứu tương lai.')

H4('Hàm ý (Implications)')
P('V-NAMHS tìm thấy nhu cầu không đáp ứng tiềm năng lớn đối với dịch vụ, với dưới 10 % những người có vấn đề SKTT tiếp cận dịch vụ trong 12 tháng qua. Tuy nhiên, các phát hiện cũng nhấn mạnh cơ hội cải thiện xu hướng VTN và gia đình tìm kiếm trợ giúp cho vấn đề SKTT. Ví dụ, các phát hiện cho thấy phụ huynh và gia đình là nguồn hỗ trợ trung tâm cho VTN. Gần ba phần tư VTN (73,9 %) báo cáo nói chuyện với thành viên gia đình khi có lo lắng. Điều này phù hợp với nghiên cứu trước tại VN cũng tìm thấy người trẻ có xu hướng tìm kiếm trợ giúp cho vấn đề SKTT từ gia đình và bạn bè (Kamimura et al. 2018, Truc Thanh Thai, Ngoc Ly Ly Thi Vu and Han Hy Thi Bui 2020). Tuy nhiên, trong V-NAMHS, chỉ 5,1 % phụ huynh xác định con cần trợ giúp cho vấn đề cảm xúc và hành vi trong 12 tháng qua, dù 21,7 % VTN có vấn đề SKTT trong cùng giai đoạn. Trong nhóm đó, hơn một phần ba (37,7 %) báo cáo con đã nhận tất cả trợ giúp cần thiết, và 20,4 % không tiếp cận vì ưa thích tự giải quyết.')

P('Kết hợp dữ liệu phụ huynh và VTN cho thấy các chiến lược được thiết kế đặc biệt để cải thiện mental health literacy, giảm kỳ thị và tăng nhận thức về dịch vụ sẵn có giữa các gia đình có thể là bước khôn ngoan. Các chương trình có thể tập trung vào giúp phụ huynh tự tin hơn khi xử lý lo lắng do VTN trình bày, đồng thời giáo dục họ về chỉ báo vấn đề SKTT, nơi tìm trợ giúp, và lợi ích của việc đó cho cả VTN và gia đình. Có bằng chứng đang phát triển về thích ứng các chương trình mental health literacy cho VN (Hoang Minh Dang et al. 2018, Nga Linh La et al. 2022, Thach Tran et al. 2020). Vấn đề mental health literacy thấp và kỳ thị đáng kể xung quanh SKTT tại VN được chứng minh bởi nghiên cứu hiện có báo cáo quan điểm từ cả quần thể chung (Kamimura et al. 2018, Mai Do et al. 2014, Mckelvey et al. 1999, Quynh Chi Nguyen Thai và Thanh Huong Nguyen 2018, Truc Thanh Thai, Ngoc Ly Ly Thi Vu và Han Hy Thi Bui 2020) và chuyên gia SKTT (Hoang-Minh Dang et al. 2020), lưu ý bằng chứng cụ thể cho VTN và gia đình vẫn còn hạn chế.')

P('Mọi cải thiện về hành vi tìm kiếm trợ giúp của VTN và gia đình cũng phải được đáp ứng bởi cải thiện tương ứng về tính sẵn có của dịch vụ. Lại nữa, phát hiện V-NAMHS nhấn mạnh cơ hội cho các bước ban đầu. Ví dụ, trong số tiếp cận dịch vụ cho vấn đề cảm xúc và hành vi 12 tháng qua, hơn một nửa (56,2 %) đã làm vậy từ nhà cung cấp y tế chính thức như bác sĩ hoặc điều dưỡng. Do đó có thể có cơ hội lồng ghép sàng lọc SKTT vào dịch vụ y tế chung hiện có, đồng thời cung cấp giáo dục và đào tạo về SKTT và đường dẫn chuyển tuyến cho người hành nghề y tế chung. Ví dụ, để tạo điều kiện lồng ghép dịch vụ chăm sóc SKTT vào y tế cơ sở, WHO đã xuất bản mhGAP Intervention Guide (phiên bản 1.0 năm 2010; phiên bản 2.0 năm 2016) cho rối loạn tâm thần, thần kinh và sử dụng chất trong môi trường y tế không chuyên khoa để bác sĩ, điều dưỡng và nhân viên y tế khác cũng như người hoạch định y tế và quản lý sử dụng (WHO 2016).')

PageMarker('31')
P('Tính khả thi và hữu ích của lồng ghép dịch vụ SKTT vào môi trường chăm sóc ban đầu và cộng đồng đã được chứng minh ở một số LMICs (Keynejad, Spagnolo and Thornicroft 2022). Hướng dẫn mhGAP cung cấp nền tảng cho triển khai các chính sách và thực hành như vậy tại VN, có thể sử dụng để tiếp tục tạo điều kiện và hỗ trợ các dịch vụ mà VTN và gia đình đang tìm kiếm cho hỗ trợ SKTT. Lồng ghép thực hành SKTT vào nhà cung cấp y tế ban đầu sẽ giúp tận dụng phát hiện V-NAMHS bằng cách cung cấp cho nhà cung cấp y tế chính thức đào tạo và tài nguyên phù hợp để hỗ trợ SKTT. Cần lưu ý rằng bất kỳ nỗ lực nào như vậy đều yêu cầu công nhận rối loạn tâm thần thông qua chính sách và chiến lược phù hợp được thiết kế để hướng dẫn phân bổ nguồn lực cho SKTT VTN được cải thiện.')

# =============================================================
# COVID-19 chapter (pages 32-35)
# =============================================================
PageMarker('32')
H2('COVID-19')

H3('Tổng quan')
P('Ca COVID-19 đầu tiên được phát hiện ở VN vào tháng 1/2020. Một Uỷ ban chỉ đạo quốc gia sau đó được thành lập để quản lý phản ứng của Chính phủ đối với đại dịch COVID-19. Chính phủ khởi động các thủ tục cách ly chính thức và biện pháp nghiêm ngặt để ngăn chặn lây lan COVID-19 gồm cấm tụ tập đông người, thực thi hạn chế di chuyển và bắt buộc đeo khẩu trang nơi công cộng. Lệnh phong toả toàn quốc bắt đầu từ tháng 4/2020, người dân phải ở nhà, doanh nghiệp và trường học đóng cửa. Sau phong toả quốc gia, các đợt phong toả địa phương cấp tỉnh tiếp tục. VN chứng kiến đỉnh COVID-19 vào tháng 9/2021 với hơn 10.000 ca mới mỗi ngày. Năm 2022, với tỷ lệ tiêm chủng cao, Chính phủ nới lỏng các biện pháp phong toả và cách ly trên toàn quốc (Bộ Y tế 2022, Chính phủ 2022).')

P('Trên toàn cầu, đại dịch COVID-19 đã tác động cuộc sống của VTN thông qua thực thi các biện pháp containment (gồm đóng cửa trường học) dẫn đến cô lập xã hội, thiếu thói quen hàng ngày và căng thẳng hộ gia đình. Bằng chứng mới nổi cho thấy mối liên hệ giữa đại dịch COVID-19 và tỷ lệ lo âu, trầm cảm, stress cao hơn (Jones, Mitra and Bhuiyan 2021). Hiểu COVID-19 có thể đã ảnh hưởng SKTT VTN như thế nào là cần thiết cho chuẩn bị các đại dịch tương lai và cung cấp hỗ trợ mà người trẻ cần để giải quyết cái mà một số gọi là "thế hệ mất mát" (Hafstad and Augusti 2021). Hiện tại, thiếu dữ liệu đại diện quốc gia từ cả VTN hoặc phụ huynh về cách đại dịch COVID-19 đã tác động cuộc sống của họ trong bối cảnh SKTT và hạnh phúc. Chương này phác thảo các yếu tố COVID-19 liên quan nhất với V-NAMHS và báo cáo các phát hiện chính. Tất cả tỷ lệ và số lượng trong chương này đều được gia trọng.')

PageMarker('33')
H3('Đo lường')
P('Các câu hỏi COVID-19 tập trung vào các yếu tố dễ liên kết với SKTT nhất theo mục tiêu cốt lõi của V-NAMHS. Các câu hỏi này không nhằm đo lường toàn diện tất cả trải nghiệm trong đại dịch COVID-19. Chúng được thiết kế đặc biệt cho khảo sát, ban đầu dựa trên rà soát tài liệu liên quan rồi tham khảo với năm nhóm NAMHS quốc tế. Câu hỏi được hỏi cả phụ huynh và VTN và liên quan cụ thể đến trải nghiệm trong đại dịch.')

P('Cả phụ huynh và VTN đều được hỏi trước hết liệu họ đã nghe về COVID-19 hay chưa. Những người đã nghe được hỏi một chuỗi câu hỏi liên quan đến trải nghiệm COVID-19. Phụ huynh được hỏi về tiếp xúc trực tiếp với COVID-19, cách ly, kỳ thị, tác động kinh tế lên hộ gia đình (gồm thay đổi thu nhập và thiếu tiền cho nhu yếu phẩm), và thay đổi trong sử dụng rượu và ma tuý của bản thân. Phụ huynh cũng được hỏi về nhu cầu trợ giúp cho vấn đề cảm xúc và hành vi của con trong đại dịch, liệu đã sử dụng dịch vụ cho các vấn đề này trong đại dịch hay không, và những rào cản liên quan COVID-19 đã ngăn họ tìm trợ giúp cho con (nếu áp dụng).')

P('VTN được hỏi về việc có dừng đi học vĩnh viễn trong đại dịch không. Các em cũng được hỏi về các trải nghiệm khác trong đại dịch gồm chứng kiến bạo lực giữa người lớn trong hộ, việc sử dụng rượu và ma tuý của người lớn trong hộ, và sử dụng rượu và ma tuý của chính mình. Cuối cùng, VTN được hỏi về gia tăng các vấn đề cảm xúc và hành vi cụ thể trong đại dịch và liệu có ai để nói chuyện khi trải qua các vấn đề này không.')

H3('Phát hiện')
P('Như Bảng 17 thể hiện, 7,7 % VTN "rất đồng ý" rằng họ thường xuyên trải qua ít nhất một vấn đề cảm xúc hoặc hành vi nhiều hơn bình thường trong đại dịch COVID-19, không có khác biệt giữa nam và nữ (nếu bao gồm cả những em "đồng ý", tỷ lệ sẽ cao tới 67 %).')

MakeTable(
    ['Vấn đề', '% Nam', 'n', '% Nữ', 'n', '% Tổng', 'n'],
    [
        ('Lo âu hoặc stress nhiều hơn', '5,2', '162', '4,9', '141', '5,1', '304'),
        ('Buồn hoặc trầm cảm hơn', '3,6', '112', '3,7', '106', '3,6', '218'),
        ('Vấn đề tập trung nhiều hơn', '2,7', '85', '1,9', '55', '2,3', '140'),
        ('Cô đơn hoặc bị cô lập hơn', '2,2', '70', '1,5', '42', '1,9', '112'),
        ('Tổng (tăng ít nhất 1 vấn đề)', '8,2', '254', '7,3', '209', '7,7', '463'),
    ])
P('Bảng 17. Tỷ lệ thường xuyên trải qua vấn đề cảm xúc và hành vi nhiều hơn bình thường trong đại dịch COVID-19 ở VTN 10–17 tuổi theo giới. Gia trọng N: nam = 3.119; nữ = 2.877. Nguồn: Bảng 17 báo cáo gốc, trang 33.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('Tổng 7,1 % (n = 427) phụ huynh báo cáo con cần trợ giúp cho vấn đề cảm xúc và hành vi trong đại dịch COVID-19. Trong số đó, 80,3 % (n = 343) không tiếp cận dịch vụ trong đại dịch, và hầu hết báo cáo lý do là sợ nhiễm COVID-19 (69,2 %; n = 237), lưu ý phụ huynh có thể chọn nhiều lý do không tiếp cận dịch vụ.')

PageMarker('34')
AddImg('p034_img1_Image238.jpg',
       'Hình 4. Trải nghiệm trong đại dịch COVID-19 ở VTN 10–17 tuổi',
       'Figure 4. Experiences during the COVID-19 pandemic among 10–17-year-olds',
       width_cm=11.0)

P('Hình 4 thể hiện tỷ lệ VTN ghi nhận các trải nghiệm khác nhau trong đại dịch COVID-19 và tỷ lệ báo cáo đây là gia tăng so với trước đại dịch. Dù gần 10 % VTN báo cáo có người lớn trong hộ uống rượu hoặc sử dụng ma tuý bất hợp pháp trong đại dịch, dưới 1 % (0,5 %) toàn mẫu VTN báo cáo đây là gia tăng so với trước đại dịch. Hơn hai phần ba phụ huynh (71,5 %) báo cáo giảm thu nhập hộ gia đình trong đại dịch COVID-19. Thêm nữa, 12,3 % nêu rằng họ thường xuyên không có đủ tiền cho nhu yếu phẩm trong đại dịch.')

H3('Suy xét (Considerations)')
H4('Diễn giải (Interpretation)')
P('Một trong mười ba VTN (7,7 %) báo cáo thường xuyên cảm thấy trầm cảm hơn, lo âu hơn, cô lập hơn, hoặc có nhiều vấn đề tập trung hơn trong đại dịch COVID-19 so với trước đại dịch. Các phát hiện này nhìn chung nhất quán với các nghiên cứu khác tại VN. Ví dụ, đánh giá nhanh định tính của UNICEF về tác động COVID-19 tại VN tìm thấy liên hệ giữa COVID-19 và tăng mức stress, lo âu, trầm cảm được người trẻ báo cáo giai thoại (UNICEF 2020a). Tuy nhiên, so sánh khó khăn vì có ít nghiên cứu định lượng về tác động COVID-19 tại VN, không có nghiên cứu nào ở quy mô lớn. Thêm nữa, dù 7,1 % phụ huynh báo cáo con cần trợ giúp cho vấn đề cảm xúc và hành vi trong đại dịch, 80,3 % nói họ không sử dụng dịch vụ. Điều này chủ yếu do sợ nhiễm COVID-19 (69,2 %). Các nghiên cứu tương tự tại LMICs khác cũng báo cáo một tỷ lệ đáng kể VTN và người trẻ cảm thấy cần trợ giúp liên quan đến hạnh phúc thể chất và tâm thần trong đại dịch COVID-19 nhưng không yêu cầu trợ giúp (UNICEF 2020b).')

H4('Hạn chế (Limitations)')
P('Dù V-NAMHS cung cấp cơ hội độc đáo đánh giá tác động của đại dịch COVID-19 lên SKTT và hạnh phúc VTN, nghiên cứu không được thiết kế để đo toàn diện tất cả các khía cạnh của đại dịch COVID-19. Các câu hỏi được thiết kế để cung cấp một "chụp nhanh" ngắn về đại dịch COVID-19 trong bối cảnh V-NAMHS và cung cấp cơ hội cho phân tích tương lai có tính đến bất kỳ tác động nào của đại dịch. Do đó, vẫn có thể một số khía cạnh liên quan của giai đoạn đại dịch không được nắm bắt. Thêm nữa, có những lo ngại ban đầu rằng nỗi sợ tiếp xúc COVID-19 có thể ảnh hưởng tỷ lệ trả lời V-NAMHS và có thể làm lệch các phát hiện. Tuy nhiên, tỷ lệ trả lời cuối cùng 81,1 % cho thấy đây không phải vấn đề.')

PageMarker('35')
H4('Hàm ý (Implications)')
P('Các phát hiện V-NAMHS chứng minh đại dịch COVID-19 đã có tác động tiêu cực lên SKTT VTN tại VN. Các phát hiện này, kết hợp với các phát hiện chỉ ra nhu cầu không được đáp ứng đối với dịch vụ SKTT, cho thấy cần cải thiện chính sách và dịch vụ hỗ trợ VTN hiện tại, tương lai, và trong chuẩn bị cho các đại dịch/tình huống khủng hoảng tương lai. Ví dụ, các dịch vụ dễ tiếp cận (ví dụ hotline điện thoại hoặc dịch vụ chat online) có thể là cách thận trọng để giải quyết nhu cầu không đáp ứng tiềm năng lớn, đồng thời bảo đảm chăm sóc khủng hoảng và liên tục trong chăm sóc SKTT hiện có bất chấp hoàn cảnh bất thường. Trong những tình huống này, nâng cao sức khoẻ có mục tiêu có thể hữu ích để nhắm tới VTN phát triển vấn đề SKTT liên quan đến các stressor cụ thể — ví dụ cô đơn và trầm cảm liên quan đóng cửa trường học. VTN cần tiếp cận dịch vụ lần đầu có thể đặc biệt dễ tổn thương, và truyền thông y tế công cộng cung cấp kiến thức, giảm kỳ thị và bình thường hoá SKTT có thể đặc biệt quan trọng. Thêm nữa, gia tăng báo cáo vấn đề cảm xúc và hành vi trong đại dịch COVID-19 nhấn mạnh tầm quan trọng của việc bao gồm SKTT trong lập kế hoạch cho các sự kiện cấp dân số trong tương lai như đại dịch, thảm hoạ tự nhiên và xung đột.')

doc.save(OUT)
print(f'Part 4 (Service Use + COVID-19) saved.')
