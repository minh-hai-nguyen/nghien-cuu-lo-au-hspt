# -*- coding: utf-8 -*-
"""Bổ sung bản dịch QT31-34 — viết lại đầy đủ hơn"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

# =====================================================================
# QT31 — 59 Countries Anxiety — Islam et al. 2025
# =====================================================================
def dich_QT31():
    print('QT31 59Countries...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.06.001', size=10)
    add_heading(doc, 'Tỷ lệ và các yếu tố liên quan của lo âu ở thanh thiếu niên đi học: Phân tích từ 59 quốc gia', 1)
    h = doc.add_paragraph(); r = h.add_run('Prevalence of and Factors Associated with Anxiety among School Going Adolescents: Analysis from 59 Countries'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Prevalence of and Factors Associated with Anxiety among School Going Adolescents: Analysis from 59 Countries'),
        ('Tác giả', 'Md. Amirul Islam, Tanjirul Islam, Bristi Rani Saha, Sakib Al Hassan, Noman Hasan, Md. Ashfikur Rahman'),
        ('Cơ quan', "Xi'an Jiao Tong University; Khulna University, Bangladesh; City University of Hong Kong"),
        ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
        ('Xuất bản', '2025, Vol. 393, 120315, 11 trang'),
        ('DOI', '10.1016/j.jad.2025.06.001'),
        ('Loại NC', 'Phân tích thứ cấp — GSHS (Global School-based Student Health Survey, WHO)'),
        ('Mẫu', '179.937 VTN 11–17 tuổi đi học, từ 59 quốc gia (chủ yếu LMIC)'),
    ])
    add_page_ref(doc, '1–11', 'J Affective Disorders', 'Vol. 393, 2025')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Bối cảnh: Tuổi VTN là giai đoạn quan trọng với thay đổi cảm xúc/tâm lý đáng kể, thường đi kèm tăng tính dễ bị tổn thương với lo âu. NC nhằm điều tra tỷ lệ lo âu và yếu tố liên quan ở VTN đi học trên 59 quốc gia.')
    p = doc.add_paragraph()
    r = p.add_run('Phương pháp: Phân tích dữ liệu GSHS (WHO), 179.937 VTN 11–17 tuổi từ 59 quốc gia. Lo âu được đánh giá dựa trên câu hỏi GSHS chuẩn hóa. Hồi quy logistic đa biến (AOR) phân tầng theo khu vực WHO.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Tỷ lệ lo âu cao ở VTN đi học toàn cầu. Yếu tố nguy cơ mạnh nhất: ý tưởng tự tử (AOR = 2,84), bất an thực phẩm nặng (AOR = 2,22), bắt nạt (AOR = 1,68), chấn thương nghiêm trọng (AOR = 1,61), ngồi >4h/ngày (AOR = 1,50), nữ (AOR = 1,51), tuổi 17 (AOR = 2,45). Yếu tố bảo vệ: cha mẹ kiểm tra bài (AOR = 0,75).')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['J Affective Disorders Q1 IF ≈ 6,6.','MẪU CỰC LỚN: 179.937 VTN từ 59 quốc gia LMIC — lớn nhất về lo âu VTN.','GSHS (WHO) — dữ liệu chuẩn hóa, đại diện, so sánh được.','Bất an thực phẩm — yếu tố ít được NC trước đó, AOR = 2,22.','Phân tầng theo 6 khu vực WHO — phát hiện khác biệt vùng.','Bao gồm nhiều nước ASEAN/ĐNA — gần bối cảnh VN.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Cắt ngang — không nhân quả.','Lo âu đo bằng 1 câu GSHS — không phải thang chuẩn (GAD-7, DASS).','Chỉ VTN đi học — loại bỏ thất nghiệp, bỏ học.','Tự báo cáo.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Lo âu là rối loạn tâm thần phổ biến nhất ở VTN: 4,4% trẻ 10–14 tuổi, 5,5% trẻ 15–19 tuổi (WHO, 2024). Tỷ lệ cao nhất ở Mỹ Latinh, Caribbean, Bắc Mỹ thu nhập cao, Tây Âu. Thấp nhất ở Nam Á và châu Phi hạ Sahara. Lo âu VTN liên quan đến suy giảm chức năng, học kém, đau khổ cảm xúc, giảm chất lượng sống, tăng nguy cơ tự tử.')
    add_p(doc, 'GSHS là khảo sát sức khỏe học đường toàn cầu do WHO và CDC phối hợp, thực hiện tại >100 quốc gia. Thiết kế lấy mẫu cụm 2 bậc, đại diện quốc gia cho VTN đi học 13–17 tuổi.')

    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Dữ liệu: GSHS (WHO/CDC), 179.937 VTN 11–17 tuổi từ 59 quốc gia. Thiết kế: cắt ngang, lấy mẫu cụm 2 bậc. Biến phụ thuộc: lo âu (câu hỏi GSHS: "Trong 12 tháng qua, bạn có thường xuyên lo lắng đến mức không thể ngủ không?"). Biến độc lập: tuổi, giới, bất an thực phẩm, bắt nạt, đánh nhau, chấn thương, bạn bè, cha mẹ, ngồi nhiều, BMI, ý tưởng/kế hoạch/cố gắng tự tử.')
    add_p(doc, 'Phân tích: Hồi quy logistic đa biến (AOR, KTC 95%). Phân tầng theo 6 khu vực WHO: Châu Phi, Đông Địa Trung Hải, Đông Nam Á, châu Mỹ, Tây Thái Bình Dương, châu Âu.')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Yếu tố liên quan lo âu ở VTN đi học (59 quốc gia, n = 179.937)', 3)
    add_table(doc,
        ['Yếu tố', 'AOR', 'KTC 95%', 'p', 'Ý nghĩa'],
        [['Tuổi 17 (vs 11–12)', '2,45', '2,05–2,94', '<0,001', 'Tăng gấp 2,5 lần'],
         ['Nữ (vs nam)', '1,51', '1,38–1,66', '<0,001', 'Nữ > Nam'],
         ['Bất an thực phẩm (nặng)', '2,22', '1,94–2,55', '<0,001', 'Yếu tố mạnh thứ 2'],
         ['Bắt nạt', '1,68', '1,52–1,85', '<0,001', ''],
         ['Chấn thương nghiêm trọng', '1,61', '1,46–1,79', '<0,001', ''],
         ['Đánh nhau', '1,26', '1,46–1,79', '<0,001', ''],
         ['Không có bạn thân', '1,28', '1,10–1,48', '0,001', ''],
         ['Cha mẹ không hiểu', '1,44', '1,27–1,63', '<0,001', ''],
         ['Cha mẹ kiểm tra bài', '0,75', '0,67–0,84', '<0,001', 'BẢO VỆ'],
         ['Ngồi >4h/ngày', '1,50', '1,32–1,71', '<0,001', ''],
         ['Ý tưởng tự tử', '2,84', '1,85–2,36', '<0,001', 'Yếu tố mạnh nhất'],
         ['Kế hoạch tự tử', '1,36', '1,21–1,54', '<0,001', ''],
         ['Cố gắng tự tử', '1,41', '1,25–1,60', '<0,001', '']],
        widths=[4.0, 1.5, 2.5, 1.5, 3.5])

    add_heading(doc, 'Bảng 2. Khác biệt theo khu vực WHO', 3)
    add_table(doc,
        ['Khu vực', 'Tỷ lệ lo âu', 'Yếu tố nổi bật', 'So sánh'],
        [['Đông Địa Trung Hải', 'Cao nhất', 'Bất an thực phẩm mạnh nhất', ''],
         ['Châu Phi', 'Cao', 'Bắt nạt, chấn thương', ''],
         ['Đông Nam Á', 'Trung bình–Cao', 'Tuổi, nữ, thực phẩm', 'Gần bối cảnh VN'],
         ['Châu Mỹ', 'Trung bình', 'Ý tưởng tự tử mạnh', ''],
         ['Tây Thái Bình Dương', 'Trung bình', 'Béo phì nổi bật hơn', 'Bao gồm TQ, Philippines'],
         ['Châu Âu', 'Thấp hơn', '', 'Ít nước tham gia GSHS']],
        widths=[3.5, 2.5, 4.0, 3.0])

    add_heading(doc, '4. THẢO LUẬN', 2)
    add_p(doc, 'Đây là đánh giá toàn diện nhất về lo âu VTN đi học (59 quốc gia, 179.937 VTN). Bất an thực phẩm (AOR = 2,22) — yếu tố ít được NC trước đó nhưng rất mạnh. Cơ chế: thiếu dinh dưỡng → stress gia đình → lo âu. Bắt nạt (AOR = 1,68) phù hợp với NC trước. Nữ > nam (AOR = 1,51) — do bạo lực giới, phân biệt đối xử, và yếu tố sinh học (estrogen).')
    add_p(doc, 'Cha mẹ kiểm tra bài (AOR = 0,75) — yếu tố bảo vệ quan trọng. Phù hợp Wen 2020 (hỗ trợ SKTT trường OR = 0,562): sự tham gia của cha mẹ và trường giảm lo âu.')
    add_p(doc, 'Ý tưởng tự tử có AOR cao nhất (2,84) — lo âu và tự tử liên quan mạnh, cần sàng lọc kết hợp.')

    add_heading(doc, '5. KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu GSHS 179.937 VTN từ 59 quốc gia cho thấy lo âu VTN phổ biến toàn cầu, liên quan mạnh với bất an thực phẩm (AOR = 2,22), bắt nạt (1,68), ý tưởng tự tử (2,84), và nữ giới (1,51). Cha mẹ tham gia (AOR = 0,75) là yếu tố bảo vệ. Cần can thiệp đa tầng: an ninh lương thực, chống bắt nạt, tăng hỗ trợ gia đình, sàng lọc tự tử.')

    add_abbreviation_table(doc, [('GSHS','Global School-based Student Health Survey — Khảo sát Sức khỏe Học đường Toàn cầu'),('AOR','Adjusted Odds Ratio — Tỷ số chênh điều chỉnh'),('LMIC','Low and Middle-Income Countries — Nước thu nhập thấp và trung bình'),('BMI','Body Mass Index — Chỉ số Khối Cơ thể'),('WHO','World Health Organization'),('CDC','Centers for Disease Control and Prevention')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['J Affective Disorders Q1. Mẫu cực lớn 179.937 VTN, 59 quốc gia.','GSHS chuẩn hóa — so sánh được giữa các nước.','Bất an thực phẩm — đóng góp mới, ít NC trước đó.','Phân tầng 6 khu vực WHO — phát hiện khác biệt vùng.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['Lo âu chỉ đo bằng 1 câu GSHS — không phải thang chuẩn (GAD-7, DASS-21). Tỷ lệ có thể không chính xác.','Cắt ngang — không xác lập nhân quả.','Chỉ VTN đi học — bỏ sót VTN không đi học (nhóm có nguy cơ cao hơn).','VN KHÔNG nằm trong 59 quốc gia này (GSHS VN chưa công bố gần đây).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN cần tham gia GSHS đầy đủ và phân tích riêng lo âu.','So sánh yếu tố nguy cơ VN vs ĐNA (Thái Lan, Philippines, Indonesia — có trong GSHS).','Bất an thực phẩm ở VN — chưa NC liên quan lo âu VTN, đặc biệt vùng nông thôn/DTTS.']:
        add_red(doc, f'• {s}')
    save(doc, '31_59Countries_Anxiety_2025.docx')

# =====================================================================
# QT32 — Ireland MyWorld — Fitzgerald et al. 2024
# =====================================================================
def dich_QT32():
    print('QT32 Ireland...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1111/eip.13530', size=10)
    add_heading(doc, 'Khám phá xu hướng thay đổi trầm cảm và lo âu ở thanh thiếu niên từ 2012 đến 2019: Từ khảo sát My World', 1)
    h = doc.add_paragraph(); r = h.add_run('Exploring Changing Trends in Depression and Anxiety among Adolescents from 2012 to 2019: Insights from My World Repeated Cross-Sectional Surveys'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Amanda Fitzgerald, Ciara Mahon, Mark Shevlin, Barbara Dooley, Aileen O. Reilly'),
        ('Cơ quan', 'UCD Dublin; Ulster University; Jigsaw Ireland'),
        ('Tạp chí', 'Early Intervention in Psychiatry (Q2, IF ≈ 2,5)'),
        ('Xuất bản', '2024, 11 trang'),
        ('DOI', '10.1111/eip.13530'),
        ('Loại NC', 'Cắt ngang lặp lại — My World Survey (2012 vs 2019)'),
        ('Mẫu', 'VTN Ireland 12–19 tuổi: Đợt 1 (2012, n = 6.085) và Đợt 2 (2019, n = 5.869)'),
    ])
    add_page_ref(doc, '1–11', 'Early Intervention in Psychiatry', '2024')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Trầm cảm (DASS-21) và lo âu (DASS-21) TĂNG đáng kể từ 2012 đến 2019 ở VTN Ireland. Nữ tăng nhanh hơn nam — đặc biệt nữ giới, tự trọng thấp, và khả năng phục hồi thấp liên quan MẠNHhơn với lo âu ở Đợt 2 so với Đợt 1. Xu hướng TRƯỚC COVID — xác nhận xu hướng tăng không chỉ do đại dịch.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['My World Survey — khảo sát SKTT thanh niên lớn nhất Ireland.','2 đợt (2012 vs 2019) — xu hướng trước COVID.','Mẫu lớn: 6.085 + 5.869 = 11.954 VTN.','Multi-group path analysis — kiểm tra yếu tố nào thay đổi giữa 2 đợt.','Yếu tố bảo vệ: lạc quan, kết nối bạn bè, tự trọng, One Good Adult (OGA).']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Chỉ Ireland — giàu, Tây Âu.','DASS-21 (không GAD-7).','Cắt ngang lặp lại — không cùng cá nhân.','Đợt 2 chỉ đến 2019 — không phản ánh hậu COVID.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. PHƯƠNG PHÁP', 2)
    add_p(doc, 'My World Survey: khảo sát cộng đồng lớn nhất Ireland về SKTT thanh niên. Đợt 1 (2012, Dooley & Fitzgerald, 2013): n = 6.085. Đợt 2 (2019, Dooley et al., 2019): n = 5.869. VTN 12–19 tuổi, khảo sát tại trường (giấy hoặc trực tuyến).')
    add_p(doc, 'Thang đo: DASS-21 (lo âu + trầm cảm), LOT-R (lạc quan), Rosenberg Self-Esteem Scale, READ (resilience), BMSLSS (hài lòng cuộc sống), MAC (kết nối), MSPSS (hỗ trợ xã hội), One Good Adult (OGA).')
    add_p(doc, 'Phân tích: Multi-group path analysis — nhóm = Đợt 1/2. So sánh hệ số hồi quy giữa 2 đợt để xác định yếu tố nào thay đổi. SPSS + Mplus.')

    add_heading(doc, '2. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xu hướng lo âu và trầm cảm 2012 → 2019', 3)
    add_table(doc,
        ['Chỉ số', 'Đợt 1 (2012)', 'Đợt 2 (2019)', 'Thay đổi', 'p'],
        [['Lo âu (DASS-21 TB)', 'Thấp hơn', 'Cao hơn', '↑ TĂNG', '<0,01'],
         ['Trầm cảm (DASS-21 TB)', 'Thấp hơn', 'Cao hơn', '↑ TĂNG', '<0,01'],
         ['Nữ — lo âu', 'Thấp hơn', 'Cao hơn', '↑ Tăng NHANH hơn nam', '<0,01'],
         ['Tự trọng → lo âu', 'Liên quan yếu hơn', 'Liên quan MẠNH hơn', '↑ Mạnh hơn ở Đợt 2', '<0,05'],
         ['Resilience → lo âu', 'Liên quan yếu hơn', 'Liên quan MẠNH hơn', '↑ Mạnh hơn ở Đợt 2', '<0,05'],
         ['Rượu → lo âu', 'Liên quan mạnh', 'Liên quan YẾU hơn', '↓ Giảm ở Đợt 2', '<0,05']],
        widths=[3.5, 2.5, 2.5, 3.0, 1.5])

    add_heading(doc, 'Bảng 2. Yếu tố bảo vệ và nguy cơ (path analysis)', 3)
    add_table(doc,
        ['Yếu tố', 'Lo âu', 'Trầm cảm', 'Thay đổi 2012→2019'],
        [['Lạc quan', 'Giảm (−)', 'Giảm (−)', 'Ổn định'],
         ['Kết nối bạn bè', 'Giảm (−)', 'Giảm (−)', 'Ổn định'],
         ['Tự trọng', 'Giảm (−)', 'Giảm (−)', 'MẠNH hơn ở 2019'],
         ['Resilience (cá nhân)', 'Giảm (−)', 'Giảm (−)', 'MẠNH hơn ở 2019'],
         ['One Good Adult (OGA)', 'Giảm (−)', 'N/A', 'Ổn định'],
         ['Nữ giới', 'Tăng (+)', 'Tăng (+)', 'MẠNH hơn ở 2019'],
         ['Rượu', 'Tăng (+)', 'Tăng (+)', 'Yếu hơn ở 2019']],
        widths=[3.5, 2.5, 2.5, 4.0])

    add_heading(doc, '3. THẢO LUẬN VÀ KẾT LUẬN', 2)
    add_p(doc, 'Lo âu và trầm cảm VTN Ireland tăng 2012→2019 (trước COVID). Phù hợp xu hướng toàn cầu: Norway 2025 (13 năm tăng), JAACAP 2024 (Mỹ lo âu gấp đôi), GBD 2025 (AAPC 0,84%). Đáng chú ý: nữ giới, tự trọng thấp, resilience thấp liên quan MẠNH HƠN với lo âu ở 2019 so với 2012 — gợi ý yếu tố này trở nên quan trọng hơn theo thời gian.')
    add_p(doc, 'One Good Adult (OGA — một người lớn tốt) là yếu tố bảo vệ — phù hợp văn hóa VN (thầy/cô, cha mẹ). Can thiệp nhắm vào xây dựng mối quan hệ mentor tại trường.')
    add_p(doc, 'Rượu giảm vai trò ở 2019 — có thể do giảm uống rượu ở VTN Ireland (xu hướng toàn cầu).')

    add_abbreviation_table(doc, [('DASS-21','Depression Anxiety Stress Scale 21'),('LOT-R','Life Orientation Test-Revised'),('READ','Resilience Scale for Adolescence'),('BMSLSS','Brief Multidimensional Students Life Satisfaction Scale'),('MAC','Hemingway Measure of Adolescent Connectedness'),('MSPSS','Multidimensional Scale of Perceived Social Support'),('OGA','One Good Adult — Một Người Lớn Tốt')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: My World lớn nhất Ireland. 2 đợt, n = 11.954. Multi-group path analysis. Trước COVID.')
    add_red(doc, 'Hạn chế: Chỉ Ireland. DASS-21 (không GAD-7). Cắt ngang lặp lại. 2019 — thiếu hậu COVID.')
    add_red(doc, 'Gap: VN cần khảo sát lặp lại tương tự (pre/post COVID). Áp dụng OGA cho VN — mentor tại trường.')
    save(doc, '32_Ireland_MyWorld_2024.docx')

# =====================================================================
# QT33 — JAMA Screen Media RCT
# =====================================================================
def dich_QT33():
    print('QT33 JAMA Screen...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1001/jamanetworkopen.2024.25438', size=10)
    add_heading(doc, 'Sử dụng phương tiện màn hình và sức khỏe tâm thần của trẻ em và thanh thiếu niên: Phân tích thứ cấp thử nghiệm lâm sàng ngẫu nhiên', 1)
    h = doc.add_paragraph(); r = h.add_run('Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Jesper Schmidt-Persson, Martin G.B. Rasmussen, et al.'),
        ('Cơ quan', 'University of Southern Denmark; University of Cambridge'),
        ('Tạp chí', 'JAMA Network Open (Q1, IF ≈ 13,8)'),
        ('Xuất bản', '2024, 12 trang'),
        ('DOI', '10.1001/jamanetworkopen.2024.25438'),
        ('Loại NC', 'Phân tích thứ cấp RCT — can thiệp giảm screen time 2 tuần'),
        ('Mẫu', '89 gia đình (181 trẻ 4–17 tuổi), Đan Mạch. 45 gia đình can thiệp, 44 đối chứng.'),
    ])
    add_page_ref(doc, '1–12', 'JAMA Network Open', '2024')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Tầm quan trọng: Sử dụng màn hình quá mức liên quan SKTT kém ở VTN qua nhiều NC quan sát. Tuy nhiên, bằng chứng thực nghiệm (RCT) hỗ trợ giả thuyết này còn THIẾU. Mục tiêu: Đánh giá tác động can thiệp giảm screen time giải trí 2 tuần lên SKTT trẻ em/VTN.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Giảm screen time → CẢI THIỆN SKTT đáng kể. Tổng khó khăn SDQ: chênh lệch −1,67 (KTC 95%: −2,68 đến −0,67), Cohen d = 0,53 (trung bình). Tác động mạnh nhất trên triệu chứng NỘI HÓA (cảm xúc + bạn bè): −1,03 (KTC: −1,76 đến −0,29). Hành vi xã hội tích cực cũng cải thiện: 0,84 (KTC: 0,39–1,30). Tuân thủ can thiệp: 97%.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['JAMA Network Open Q1 IF ≈ 13,8.','RCT — bằng chứng NHÂN QUẢ (hiếm hoi trong lĩnh vực screen time).','Cohen d = 0,53 (trung bình) — kích thước hiệu ứng có ý nghĩa lâm sàng.','Tác động mạnh nhất trên NỘI HÓA (cảm xúc) — phù hợp lo âu.','Tuân thủ 97% — can thiệp khả thi.','Đo khách quan: tracker smartphone + tablet + PC + TV monitor.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Mẫu NHỎ: 89 gia đình (181 trẻ). Can thiệp CHỈ 2 tuần.','Chỉ Đan Mạch — giàu, Bắc Âu.','Không mù (gia đình biết nhóm) — thiên lệch kỳ vọng.','Phân tích thứ cấp — không phải kết quả chính RCT gốc.','SDQ đo bởi phụ huynh — có thể thiên lệch.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. PHƯƠNG PHÁP', 2)
    add_p(doc, 'RCT 2 nhóm: can thiệp (45 gia đình, 86 trẻ) vs đối chứng (44 gia đình, 95 trẻ). Can thiệp: giảm screen time giải trí xuống ≤3 giờ/tuần/gia đình (từ mức thông thường ~17–20 giờ/tuần) trong 2 tuần. Đo: SDQ (Strengths and Difficulties Questionnaire) — phụ huynh báo cáo. Baseline: 6 tháng trước. Follow-up: 2 tuần sau. Tuân thủ đo khách quan: tracker app + PC software + TV monitor.')
    add_p(doc, 'Phân tích: Mixed-effects tobit regression. Stata 18. ITT (intention-to-treat).')

    add_heading(doc, '2. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Tác động can thiệp giảm screen time lên SKTT', 3)
    add_table(doc,
        ['Chỉ số SDQ', 'Chênh lệch (can thiệp − đối chứng)', 'KTC 95%', 'Cohen d', 'Ý nghĩa'],
        [['Tổng khó khăn', '−1,67', '−2,68 đến −0,67', '0,53', 'CẢI THIỆN — trung bình'],
         ['Triệu chứng nội hóa', '−1,03', '−1,76 đến −0,29', '—', 'CẢI THIỆN — mạnh nhất'],
         ['Triệu chứng ngoại hóa', '−0,55', '−1,20 đến 0,10', '—', 'Không đáng kể'],
         ['Hành vi xã hội tích cực', '0,84', '0,39 đến 1,30', '—', 'CẢI THIỆN']],
        widths=[3.5, 4.0, 3.0, 1.5, 3.0])

    add_heading(doc, 'Bảng 2. So sánh với NC screen time khác trong Đề tài', 3)
    add_table(doc,
        ['Nghiên cứu', 'Thiết kế', 'Phát hiện', 'Kết luận'],
        [['JAMA 2024 (bài này)', 'RCT 2 tuần', 'Giảm ST → cải thiện SKTT (d=0,53)', 'Bằng chứng nhân quả'],
         ['Li 2025 (QT22)', 'Dọc 12 tháng', 'Liên quan dọc YẾU', 'Quan hệ HAI CHIỀU?'],
         ['Norway 2025 (QT21)', 'Decomposition 13 năm', 'MXH giải thích xu hướng tăng', 'Đóng góp một phần'],
         ['Nature 2025 (QT27)', 'Cắt ngang + chẩn đoán', 'VTN SKTT dùng MXH nhiều hơn', 'g = 0,46']],
        widths=[3.0, 2.5, 4.0, 3.5])

    add_heading(doc, '3. KẾT LUẬN', 2)
    add_p(doc, 'RCT cho thấy giảm screen time giải trí CẢI THIỆN SKTT trẻ em — bằng chứng nhân quả hiếm hoi (Cohen d = 0,53). Tác động mạnh nhất trên triệu chứng NỘI HÓA (cảm xúc/bạn bè) — phù hợp lo âu. Kết hợp với Li 2025 (dọc yếu) và Norway 2025 (MXH đóng góp xu hướng), gợi ý: tác động THỰC nhưng cần GIẢM ĐỦ MẠNH (>80% thời gian) để thấy hiệu quả.')

    add_abbreviation_table(doc, [('SDQ','Strengths and Difficulties Questionnaire — Bảng hỏi Điểm mạnh và Khó khăn'),('RCT','Randomized Controlled Trial'),('ITT','Intention-to-Treat — Phân tích Ý định Điều trị'),('ST','Screen Time — Thời gian Sử dụng Màn hình')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: JAMA Q1. RCT — bằng chứng cao nhất. Đo khách quan. Cohen d = 0,53.')
    add_red(doc, 'Hạn chế: Mẫu rất nhỏ (89 gia đình). 2 tuần. Không mù. Đan Mạch. Phân tích thứ cấp. SDQ phụ huynh.')
    add_red(doc, 'Gap: VN cần RCT giảm screen time tại trường — CHƯA CÓ. Kết hợp giảm ST + CBT? Can thiệp toàn gia đình ở VN?')
    save(doc, '33_JAMA_ScreenMedia_2024.docx')

# =====================================================================
# QT34 — Korea MH Trends
# =====================================================================
def dich_QT34():
    print('QT34 Korea...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1038/s41598-024-XXXXX', size=10)
    add_heading(doc, 'Xu hướng quốc gia về sức khỏe tâm thần thanh thiếu niên theo mức thu nhập tại Hàn Quốc, trước và sau COVID-19, 2006–2022', 1)
    h = doc.add_paragraph(); r = h.add_run("National Trends in Adolescents' Mental Health by Income Level in South Korea, Pre- and Post-COVID-19, 2006-2022"); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Jaehyeong Cho, Jaeyu Park, et al.'),
        ('Tạp chí', 'Nature Scientific Reports (Q1, IF ≈ 4,6)'),
        ('Xuất bản', '2024, 14 trang, 4 bảng, 1 hình'),
        ('Loại NC', 'Phân tích xu hướng 16 năm — KYRBS quốc gia'),
        ('Mẫu', 'KYRBS (Korea Youth Risk Behavior Web-based Survey), THCS + THPT, 2006–2022'),
    ])
    add_page_ref(doc, '1–14', 'Nature Scientific Reports', '2024')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('NC đầu tiên theo dõi dài hạn 17 năm mối quan hệ thu nhập hộ gia đình — SKTT VTN, cùng xu hướng. Phát hiện: thu nhập thấp liên quan với stress, buồn bã, ý tưởng tự tử, cố tự tử. COVID-19 thay đổi đáng kể xu hướng: mô hình GIẢM trước COVID → TĂNG sau COVID. Bất bình đẳng thu nhập MỞ RỘNG trong và sau COVID — nhóm nghèo bị ảnh hưởng nặng hơn.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['Nature Sci Rep Q1 IF = 4,6.','KYRBS quốc gia 16 năm — xu hướng dài nhất trong Đề tài.','Phân tầng theo THU NHẬP — ít NC nào làm.','Mô hình ĐẢO CHIỀU: cải thiện trước COVID → xấu đi sau — KHÁC phương Tây (tăng liên tục).','Bất bình đẳng MỞ RỘNG: stress nhóm nghèo 62,8% vs giàu 40,1% — chênh 22,7 điểm.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Chỉ Hàn Quốc — văn hóa khác VN.','Đo stress/buồn bã chung — không tách lo âu riêng (không GAD-7).','KYRBS tự báo cáo trực tuyến.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. PHƯƠNG PHÁP', 2)
    add_p(doc, 'KYRBS: Khảo sát Hành vi Nguy cơ Thanh niên Hàn Quốc, triển khai hàng năm bởi KCDC và Bộ Giáo dục từ 2005, lấy mẫu cụm nhiều bậc, đại diện quốc gia cho HS THCS + THPT. Phân tầng: 5 mức thu nhập gia đình (cao → thấp). Giai đoạn: trước COVID (2006–2019) vs sau COVID (2020–2022).')
    add_p(doc, 'Phân tích xu hướng: hồi quy tuyến tính, β (95% CI). Kiểm tra thay đổi xu hướng trước/sau COVID bằng βdiff.')

    add_heading(doc, '2. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xu hướng SKTT VTN Hàn Quốc 2006–2022', 3)
    add_table(doc,
        ['Chỉ số', 'Trước COVID (β)', 'Sau COVID (β)', 'Thay đổi (βdiff)', 'Bất bình đẳng'],
        [['Stress nhận thức cao', '−1,41 (giảm)', '2,39 (TĂNG)', '3,80***', 'Nghèo 62,8% vs Giàu 40,1%'],
         ['Buồn bã', 'Giảm', 'TĂNG (28,2%)', 'Đáng kể***', 'Khoảng cách MỞ RỘNG'],
         ['Ý tưởng tự tử', 'Giảm', 'TĂNG (13,9%)', 'Đáng kể***', 'Nghèo tệ hơn'],
         ['Cố tự tử', 'Giảm', 'TĂNG', 'Đáng kể***', 'Nghèo tệ hơn']],
        widths=[3.0, 2.5, 2.5, 2.5, 4.0])

    add_heading(doc, 'Bảng 2. So sánh xu hướng Hàn Quốc với các nước', 3)
    add_table(doc,
        ['Nước', 'Giai đoạn', 'Xu hướng', 'Đặc thù'],
        [['Hàn Quốc (bài này)', '2006–2022 (16 năm)', 'Giảm trước → TĂNG sau COVID', 'ĐẢO CHIỀU'],
         ['Hoa Kỳ (QT23)', '2013–2021 (8 năm)', 'Lo âu tăng gấp đôi (AOR=2,17)', 'Tăng liên tục'],
         ['Na Uy (QT21)', '2011–2024 (13 năm)', 'Tăng liên tục', 'Trường + MXH'],
         ['Ireland (QT32)', '2012–2019 (7 năm)', 'Tăng (trước COVID)', 'Nữ nhanh hơn'],
         ['VN (VN14)', '2021 vs 2023', '41,5% → 25,4%', 'Phục hồi sau COVID?'],
         ['Toàn cầu (QT30)', '1990–2021 (31 năm)', 'AAPC 0,84%', 'GBD 204 nước']],
        widths=[3.0, 3.0, 3.5, 3.0])

    add_heading(doc, 'Bảng 3. Bất bình đẳng thu nhập', 3)
    add_table(doc,
        ['Thu nhập', 'Stress 2022', 'Buồn bã', 'Tự tử', 'So với VN'],
        [['Cao nhất', '40,1%', 'Thấp hơn', 'Thấp hơn', ''],
         ['Thấp nhất', '62,8%', 'Cao hơn', 'Cao hơn', ''],
         ['Chênh lệch', '22,7 điểm', 'MỞ RỘNG sau COVID', '', 'VN: chưa có dữ liệu']],
        widths=[2.5, 2.5, 3.0, 2.5, 3.0])

    add_heading(doc, '3. THẢO LUẬN VÀ KẾT LUẬN', 2)
    add_p(doc, 'Hàn Quốc là nước DUY NHẤT trong Đề tài cho thấy SKTT VTN CẢI THIỆN trước COVID (2006–2019) rồi XẤU ĐI sau COVID — mô hình đảo chiều, khác phương Tây (tăng liên tục). Gợi ý: (1) can thiệp SKTT CÓ THỂ hiệu quả (giai đoạn cải thiện), (2) COVID đảo ngược tiến bộ, đặc biệt ở nhóm nghèo.')
    add_p(doc, 'Bất bình đẳng thu nhập MỞ RỘNG sau COVID — stress nhóm nghèo 62,8% vs giàu 40,1% (chênh 22,7 điểm). COVID tác động KHÔNG ĐỀU — nhóm nghèo bị ảnh hưởng nặng hơn. Tại VN, chưa có dữ liệu phân tầng theo thu nhập cho SKTT VTN — GAP quan trọng.')

    add_abbreviation_table(doc, [('KYRBS','Korea Youth Risk Behavior Web-based Survey'),('KCDC','Korea Centers for Disease Control and Prevention'),('β','Beta — Hệ số xu hướng'),('βdiff','Beta difference — Chênh lệch xu hướng trước/sau COVID')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: Nature Q1. KYRBS 16 năm. Phân tầng thu nhập. Mô hình đảo chiều.')
    add_red(doc, 'Hạn chế: Chỉ Hàn Quốc. Đo stress/buồn bã chung (không GAD-7). Tự báo cáo.')
    add_red(doc, 'Gap: VN chưa có phân tích SKTT VTN theo thu nhập. GSHS VN cần phân tầng kinh tế.')
    save(doc, '34_Korea_MH_Trends_2024.docx')

# RUN ALL
dich_QT31()
dich_QT32()
dich_QT33()
dich_QT34()
print('\n=== DONE QT31-34 (bo sung day du) ===')
