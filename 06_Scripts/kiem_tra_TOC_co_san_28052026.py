# -*- coding: utf-8 -*-
"""Kiem tra TOC co san trong file - text hay field?
Doc va so sanh ten muc voi heading thuc te.
28/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

d = Document(FILE)

# 1. Find MỤC LỤC heading
toc_start = None
for i, p in enumerate(d.paragraphs[:200]):
    if p.text.strip().upper() in ('MỤC LỤC', 'MUC LUC'):
        toc_start = i
        break

print(f"MỤC LỤC at para {toc_start}")
print()

# 2. Read paragraphs after MỤC LỤC until the next H1-like heading
print("=== Existing TOC content (paras after MỤC LỤC) ===")
print()
import re
for j in range(toc_start + 1, min(toc_start + 80, len(d.paragraphs))):
    p = d.paragraphs[j]
    text = p.text
    if not text.strip():
        continue
    # Check if this is end of TOC (next major heading)
    upper = text.strip().upper()
    if any(upper.startswith(kw) for kw in ['DANH MỤC', 'MỞ ĐẦU', 'CHƯƠNG']):
        # Could be end of TOC OR a TOC entry pointing to one of these
        # If text is short (< 200) and looks like a real heading, it's end
        if len(text) < 150 and not re.search(r'\d+$', text.strip()):
            # No trailing page number → real heading, end of TOC
            print(f"  --- End of TOC at para {j}: {text!r} ---")
            break
    print(f"  [{j}] {text!r}")

# 3. Check if there are any TOC field codes
print()
print("=== TOC field codes ===")
n_toc_fields = 0
for elem in d.element.body.iter():
    if elem.tag == qn('w:instrText'):
        if elem.text and 'TOC' in elem.text:
            n_toc_fields += 1
            print(f"  TOC field instruction: {elem.text!r}")
print(f"  Total TOC fields: {n_toc_fields}")

# 4. Check if TOC entries have hyperlinks to bookmarks (field-based) or are plain text
print()
print("=== Are TOC entries clickable hyperlinks? ===")
for j in range(toc_start + 1, min(toc_start + 30, len(d.paragraphs))):
    p = d.paragraphs[j]
    if not p.text.strip(): continue
    # Check for w:hyperlink in this paragraph
    has_hl = False
    for elem in p._element.iter():
        if elem.tag == qn('w:hyperlink'):
            has_hl = True
            anchor = elem.get(qn('w:anchor'))
            print(f"  Para {j}: hyperlink anchor={anchor!r}")
            break
    if not has_hl:
        print(f"  Para {j}: no hyperlink — plain text")
    if j > toc_start + 5:
        break
