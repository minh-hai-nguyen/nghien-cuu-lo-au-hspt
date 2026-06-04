# -*- coding: utf-8 -*-
"""
KG CROSS-CHECK vs REPORT v3 — xac minh tung claim trong bao cao co trong KG.

Logic:
1. Extract claims from report v3 (pattern: "Author year" + "stat value")
2. For each claim, check if KG has matching (Paper, EffectSize) tuple
3. Report "ghost claims" (in report but not in KG) + "orphan facts" (in KG but not in report)
"""
import os, sys, io, re, json
from collections import defaultdict
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import networkx as nx
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
KG_DIR = os.path.join(os.path.dirname(__file__), 'kg_data')
REPORT = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 11042026 v3.docx')

G = nx.read_graphml(os.path.join(KG_DIR, 'kg_v1.graphml'))

# Map author name → paper IDs (rough)
author_to_papers = defaultdict(set)
for n in G.nodes:
    if G.nodes[n].get('type') == 'Author':
        name = G.nodes[n].get('name', '')
        # Find papers pointing to this author
        for u, v, d in G.in_edges(n, data=True):
            if d.get('rel') == 'AUTHORED_BY':
                author_to_papers[name.lower()].add(u)

# Also map by descriptor keywords
paper_descriptors = {}
for n in G.nodes:
    if G.nodes[n].get('type') == 'Paper':
        paper_descriptors[n] = G.nodes[n].get('descriptor', '').lower()

def find_paper_by_claim(text):
    """Given a sentence mentioning an author+year, find matching paper ID."""
    # Try common author patterns
    author_year_pat = re.compile(r'(\w[\w\-]{2,})\s+(?:et al\.\s+)?(\d{4})')
    matches = []
    for m in author_year_pat.finditer(text):
        surname = m.group(1).lower()
        year = m.group(2)
        # Look for paper with this surname + year in descriptor
        for pid, desc in paper_descriptors.items():
            if surname in desc and year in desc:
                matches.append(pid)
    return list(set(matches))

# Read report
doc = Document(REPORT)
paras = [p.text for p in doc.paragraphs if p.text.strip()]

# Extract effect size claims
claim_pat = re.compile(
    r'([^.]*?(?:Cohen\'?s?\s*d|Hedges\'?\s*g|\bd\s*=|\bg\s*=|\bOR\s*=|\bAOR\s*=|\bSMD\s*=|\bβ\s*=|\bSUCRA\s*[=:])\s*(-?\d+[,.]\d+)[^.]*)',
    re.IGNORECASE
)

ghost_claims = []  # claims in report not in KG
matched_claims = []

for para in paras:
    for m in claim_pat.finditer(para):
        full_claim = m.group(1).strip()
        val_str = m.group(2).replace(',', '.')
        try:
            val = float(val_str)
        except:
            continue
        # Find paper in this paragraph
        paper_ids = find_paper_by_claim(para)
        if not paper_ids:
            # Try looking backward at previous paragraphs
            continue
        # For each candidate paper, check if KG has matching ES
        found = False
        for pid in paper_ids:
            # Look in KG for ES nodes belonging to this paper
            for succ in G.successors(pid):
                if G.nodes[succ].get('type') == 'EffectSize':
                    kg_val = G.nodes[succ].get('value')
                    if kg_val is not None:
                        try:
                            if abs(float(kg_val) - val) < 0.01:
                                found = True
                                matched_claims.append((pid, val, full_claim[:100]))
                                break
                        except: pass
            if found: break
        if not found:
            ghost_claims.append((paper_ids, val, full_claim[:120]))

print(f'\n=== REPORT v3 CLAIM CROSS-CHECK ===')
print(f'Matched claims: {len(matched_claims)}')
print(f'Ghost claims (in report but not in KG): {len(ghost_claims)}')

print('\n--- GHOST CLAIMS (first 15) ---')
for papers, val, claim in ghost_claims[:15]:
    print(f'  {papers} val={val}')
    print(f'    "{claim}"')

# Find orphan facts (in KG but not in report text)
report_text = '\n'.join(paras)
for tb in doc.tables:
    for row in tb.rows:
        for cell in row.cells:
            report_text += '\n' + cell.text

orphan_facts = []
for n in G.nodes:
    if G.nodes[n].get('type') == 'EffectSize':
        val = G.nodes[n].get('value')
        paper = G.nodes[n].get('paper', '')
        if val is None: continue
        val_str_vn = f'{val:.3f}'.rstrip('0').rstrip('.').replace('.', ',')
        val_str_2 = f'{val:.2f}'.replace('.', ',')
        if val_str_vn not in report_text and val_str_2 not in report_text:
            orphan_facts.append((paper, G.nodes[n].get('es_type', ''), val))

print(f'\n--- ORPHAN FACTS (in KG but not in report v3) ---')
print(f'Total: {len(orphan_facts)}')
print('Top 15:')
for p, t, v in orphan_facts[:15]:
    desc = G.nodes[p].get('descriptor', '') if p in G.nodes else ''
    print(f'  {p} {t}={v} ({desc[:40]})')

# Save
with open(os.path.join(KG_DIR, 'crosscheck_v3.json'), 'w', encoding='utf-8') as f:
    json.dump({
        'matched_claims_count': len(matched_claims),
        'ghost_claims_count': len(ghost_claims),
        'orphan_facts_count': len(orphan_facts),
        'ghost_claims': [{'papers': p, 'value': v, 'claim': c}
                         for p, v, c in ghost_claims],
        'orphan_facts': [{'paper': p, 'type': t, 'value': v}
                         for p, t, v in orphan_facts],
    }, f, ensure_ascii=False, indent=2)
print(f'\nSaved: kg_data/crosscheck_v3.json')
