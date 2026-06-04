# -*- coding: utf-8 -*-
"""
QA automated: so sanh tom tat vs PDF goc — tim ra loi.
Rut key facts tu PDF + so sanh voi tom tat.
"""
import sys, os, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import fitz
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')

def read_docx(path):
    if not os.path.exists(path):
        return ''
    d = Document(path)
    txt = '\n'.join(p.text for p in d.paragraphs)
    for tb in d.tables:
        for r in tb.rows:
            for c in r.cells:
                txt += '\n' + c.text
    return txt

def read_pdf(path):
    full = os.path.join(ROOT, path) if not os.path.isabs(path) else path
    if not os.path.exists(full):
        return ''
    try:
        d = fitz.open(full)
        return '\n'.join(d[i].get_text() for i in range(len(d)))
    except Exception as e:
        print(f'  ERR reading {full}: {e}')
        return ''

def check_fact(summary, pdf_text, label, pattern, must_be_in_pdf=True):
    """Check if summary contains a claim; if so, verify it appears in pdf text"""
    # Find claim in summary
    m = re.search(pattern, summary)
    if not m:
        return None  # Not claimed
    claim = m.group(0)
    # Check in pdf
    if must_be_in_pdf and claim in pdf_text:
        return ('OK', claim)
    return ('MISSING', claim)

# Map: summary file → (pdf path, list of key facts to verify)
QA_MAP = [
    # (tt_file, pdf_path, [list of (label, regex, expected_value_in_pdf)])
    ('VN23_NguyenLX_2023_Medicine.docx',
     '02_Papers-goc/Viet-Nam/medi-102-e33559.pdf',
     [
         ('n', r'5[\.,]?730', '5,730'),
         ('GAD-7', r'GAD-?7', 'GAD-7'),
         ('prevalence', r'16[,\.]2', '16'),
         ('year', r'2023', '2023'),
         ('author', r'Nguyen,?\s*L\.?X\.?', 'Nguyen'),
         ('PR=1.71', r'PR\s*=\s*1[,.]71', '1.71'),
     ]),
    ('VN24_VinhLong_2024.docx',
     '02_Papers-goc/Viet-Nam/92-10696-18832_Văn bản của bài báo.pdf',
     [
         ('n=919', r'919', '919'),
         ('CES-D', r'CES-?D', 'CES-D'),
         ('ESSA', r'ESSA', 'ESSA'),
         ('trầm cảm 12,2', r'12[,\.]2', '12,2'),
         ('Nguyễn Thanh Truyền', r'Nguyễn Thanh Truyền', 'Truyền'),
         ('Vĩnh Long', r'Vĩnh Long', 'Vĩnh Long'),
     ]),
    ('VN25_HaiPhong_2024.docx',
     '02_Papers-goc/Viet-Nam/O2401115.pdf',
     [
         ('n=420', r'420', '420'),
         ('DASS-21', r'DASS-?21', 'DASS-21'),
         ('lo âu 39,3', r'39[,\.]3', '39,3'),
         ('Hải Phòng', r'Hải Phòng', 'Hải Phòng'),
         ('Vĩnh Bảo', r'Vĩnh Bảo', 'Vĩnh Bảo'),
     ]),
    ('VN26_LongAn_2025.docx',
     '02_Papers-goc/Viet-Nam/TranDucSi_2025_LongAn_DASS_PNT.pdf',
     [
         ('Long An', r'Long An', 'Long An'),
         ('57,2', r'57[,\.]2', '57,2'),
         ('Trần Đức Sĩ', r'Trần Đức Sĩ', 'Trần Đức Sĩ'),
     ]),
    ('VN27_Dinh_SchoolFactors_2021.docx',
     '02_Papers-goc/Viet-Nam/School_Factors_Causing_Vietnamese_Adolescents_Anx.pdf',
     [
         ('Dinh', r'Dinh', 'Dinh'),
         ('school factor', r'school.factor', 'school'),
         ('anxiety', r'anxiety', 'anxiety'),
     ]),
    ('VN28_DaoThiNgoan_2025_SVY4_HMU.docx',
     '02_Papers-goc/Viet-Nam/TCNCYH_2025_SinhVien_LoAu_TramCam.pdf',
     [
         ('n=196', r'196', '196'),
         ('33,67', r'33[,\.]67', '33,67'),
         ('43,88', r'43[,\.]88', '43,88'),
         ('31,63', r'31[,\.]63', '31,63'),
         ('4,97', r'4[,\.]97', '4,97'),
         ('Đào Thị Ngoãn', r'Đào Thị Ngoãn', 'Đào Thị Ngoãn'),
     ]),
    ('VN29_Duong_2025_2631HS_TPHCM.docx',
     '02_Papers-goc/Viet-Nam/Duong_2025_SocPsychiatry_2631HS_TPHCM.pdf',
     [
         ('n=2,631', r'2[,\.]631', '2,631'),
         ('42,6', r'42[,\.]6', '42.6'),
         ('50,3', r'50[,\.]3', '50.3'),
         ('91,6', r'91[,\.]6', '91.6'),
         ('Duong', r'Duong', 'Duong'),
     ]),
    ('QT38_DeSilva_SriLanka_RCT_2024.docx',
     '02_Papers-goc/The-gioi-moi/s13034-024-00799-9.pdf',
     [
         ('720 HS', r'720', '720'),
         ('36 trường', r'36\s*(school|trường)', '36'),
         ('De Silva', r'De Silva', 'De Silva'),
         ('Sri Lanka', r'Sri Lanka', 'Sri Lanka'),
     ]),
    ('QT39_Xian_NMA_SAD_2024.docx',
     '02_Papers-goc/The-gioi-moi/1-s2.0-S0165032724013156-main.pdf',
     [
         ('30 RCT', r'30 RCT', '30'),
         ('1.547', r'1[,\.]547', '1,547'),
         ('71,2', r'71[,\.]2', '71.2'),
         ('Xian', r'Xian', 'Xian'),
     ]),
    ('QT40_Walder_DMHI_SAD_2025.docx',
     '02_Papers-goc/The-gioi-moi/preprint-67067-accepted.pdf',
     [
         ('Walder', r'Walder', 'Walder'),
         ('0,508', r'0[,\.]508', '0.508'),
         ('0,878', r'0[,\.]878', '0.878'),
     ]),
    ('QT41_Zheng_MXH_Stress_Sleep_2025.docx',
     '02_Papers-goc/The-gioi-moi/prbm-18-1571.pdf',
     [
         ('n=469', r'469', '469'),
         ('Zheng', r'Zheng', 'Zheng'),
         ('0,615', r'0[,\.]615', '0.615'),
     ]),
    ('QT42_Brown_Carter_UK_2025.docx',
     '02_Papers-goc/The-gioi-moi/School based interventions for depression and anxiety in UK.pdf',
     [
         ('Brown', r'Brown', 'Brown'),
         ('Carter', r'Carter', 'Carter'),
         ('BESST', r'BESST', 'BESST'),
         ('8.376', r'8[,\.]?376', '8,376'),
     ]),
    ('QT43_Bress_Maya_App_2024.docx',
     '02_Papers-goc/The-gioi-moi/bress_2024_oi_240871_1723229502.27231.pdf',
     [
         ('Bress', r'Bress', 'Bress'),
         ('Maya', r'Maya', 'Maya'),
         ('NCT05130281', r'NCT05130281', 'NCT05130281'),
         ('n=59', r'59', '59'),
         ('1,04', r'1[,\.]04', '1.04'),
     ]),
    ('QT44_Cao_Resilience_MA_2025.docx',
     '02_Papers-goc/The-gioi-moi/fpsyt-16-1594658.pdf',
     [
         ('resilience', r'resilience', 'resilience'),
         ('Frontiers', r'Frontiers', 'Frontiers'),
     ]),
    ('QT45_Sasaki_Japan_iCBT_2024.docx',
     '02_Papers-goc/The-gioi-moi/pediatrics-2024-1-e55786.pdf',
     [
         ('n=77', r'77', '77'),
         ('Sasaki', r'Sasaki', 'Sasaki'),
         ('4,97', r'4[,\.]97', '4.97'),
         ('3,95', r'3[,\.]95', '3.95'),
         ('UMIN000049768', r'UMIN000049768', 'UMIN000049768'),
     ]),
    ('QT46_AcademicStress_SR_2025.docx',
     '02_Papers-goc/The-gioi-moi/s10578-024-01667-5.pdf',
     [
         ('Academic Stress', r'Academic Stress', 'Academic Stress'),
         ('Child Psychiatry', r'Child Psychiatry', 'Child Psychiatry'),
     ]),
    ('QT47_Dong_PLOS_DASS_2025.docx',
     '02_Papers-goc/The-gioi-moi/journal.pone.0328785.pdf',
     [
         ('n=2.716', r'2[,\.]716', '2,716'),
         ('24,4', r'24[,\.]4', '24.4'),
         ('41,4', r'41[,\.]4', '41.4'),
         ('15,6', r'15[,\.]6', '15.6'),
         ('0,22', r'0[,\.]22', '0.22'),
         ('Dong', r'Dong', 'Dong'),
         ("Ya'an", r"Ya.?[\'?]an", "Ya'an"),
     ]),
]

print('=' * 80)
print('QA CHECK — 22 summaries vs source PDFs')
print('=' * 80)

total_checks = 0
total_ok = 0
errors = []

for tt_file, pdf_path, facts in QA_MAP:
    tt_path = os.path.join(TT_DIR, tt_file)
    print(f'\n[{tt_file}]')
    summary = read_docx(tt_path)
    pdf_text = read_pdf(pdf_path)
    if not pdf_text:
        print(f'  ⚠ PDF not readable: {pdf_path}')
        continue
    for label, summary_pattern, expected_in_pdf in facts:
        total_checks += 1
        # Check if summary claims this
        m = re.search(summary_pattern, summary)
        in_summary = bool(m)
        # Check if pdf has the expected value
        in_pdf = expected_in_pdf.lower() in pdf_text.lower()
        if in_summary and in_pdf:
            print(f'  OK   [{label}]')
            total_ok += 1
        elif in_summary and not in_pdf:
            print(f'  ⚠ CLAIM NOT IN PDF [{label}]')
            errors.append(f'{tt_file}: "{label}" claimed in summary but not found in PDF')
        elif not in_summary and in_pdf:
            print(f'  ? FACT IN PDF BUT NOT CLAIMED [{label}]')
        else:
            print(f'  ✗ NEITHER [{label}] ')
            errors.append(f'{tt_file}: "{label}" missing from both')

print('\n' + '=' * 80)
print(f'RESULT: {total_ok}/{total_checks} facts verified OK')
print('=' * 80)

if errors:
    print(f'\n=== {len(errors)} ERRORS/WARNINGS ===')
    for e in errors:
        print(f'  • {e}')
