# -*- coding: utf-8 -*-
"""Dịch B8 — Sri Lanka CBT school-based cluster RCT 2024"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1186/s13034-024-00799-9', size=10)

add_heading(doc, 'Hiệu quả can thiệp dựa trên liệu pháp nhận thức\u2013hành vi (CBT) trong giảm lo âu ở thanh thiếu niên tại quận Colombo, Sri Lanka: Thử nghiệm ngẫu nhiên cụm có đối chứng', 1)
h = doc.add_paragraph()
r = h.add_run('Effectiveness of a cognitive behavioural therapy (CBT)-based intervention for reducing anxiety among adolescents in the Colombo District, Sri Lanka: cluster randomized controlled trial')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('Ti\u00eau \u0111\u1ec1 g\u1ed1c', 'Effectiveness of a CBT-based intervention for reducing anxiety among adolescents in the Colombo District, Sri Lanka: cluster randomized controlled trial'),
    ('T\u00e1c gi\u1ea3', 'Sinha De Silva (1)*, Renuka Peris (2), Sudharshi Senaviratne (3), Dulani Samaranayake (4)'),
    ('C\u01a1 quan', '(1) PGIM, \u0110H Colombo; (2) B\u1ed9 Gi\u00e1o d\u1ee5c (ngh\u1ec9 h\u01b0u); (3) Central Queensland Hospital, \u00dac; (4) Y khoa, \u0110H Colombo, Sri Lanka'),
    ('T\u1ea1p ch\u00ed', 'Child and Adolescent Psychiatry and Mental Health (Q1, IF \u2248 4,0)'),
    ('Xu\u1ea5t b\u1ea3n', '2024, 18:108, 12 trang'),
    ('DOI', '10.1186/s13034-024-00799-9'),
    ('Lo\u1ea1i NC', 'Th\u1eed nghi\u1ec7m ng\u1eabu nhi\u00ean c\u1ee5m c\u00f3 \u0111\u1ed1i ch\u1ee9ng (cluster RCT) \u2014 \u0111\u0103ng k\u00fd SLCTR/2018/018'),
    ('M\u1eabu', '36 tr\u01b0\u1eddng (18 can thi\u1ec7p + 18 \u0111\u1ed1i ch\u1ee9ng), 720 HS l\u1edbp 9 (~14 tu\u1ed5i), qu\u1eadn Colombo, Sri Lanka'),
    ('C\u00f4ng c\u1ee5', 'SCARED (41 m\u1ee5c, lo \u00e2u) + DASS-21 (tr\u1ea7m c\u1ea3m) + Rosenberg Self-Esteem (t\u1ef1 tr\u1ecdng)'),
    ('Can thi\u1ec7p', '8 phi\u00ean CBT/tu\u1ea7n (40 ph\u00fat) + 1 th\u00e1ng th\u1ef1c h\u00e0nh + 1 phi\u00ean cha m\u1eb9. GV \u0111\u01b0\u1ee3c \u0111\u00e0o t\u1ea1o cung c\u1ea5p.'),
    ('Truy c\u1eadp', 'Open Access \u2014 CC BY 4.0'),
])
add_page_ref(doc, '1\u201312', 'Child Adol Psychiatry Mental Health', '2024, 18:108')

# TOM TAT
add_heading(doc, 'T\u00d3M T\u1eaeT', 2)
add_p(doc, 'R\u1ed1i lo\u1ea1n lo \u00e2u l\u00e0 v\u1ea5n \u0111\u1ec1 t\u00e2m l\u00fd ph\u1ed5 bi\u1ebfn nh\u1ea5t \u1edf tr\u1ebb em v\u00e0 VTN. CBT \u0111\u01b0\u1ee3c ch\u1ee9ng minh hi\u1ec7u qu\u1ea3 gi\u1ea3m lo \u00e2u. M\u1ee5c \u0111\u00edch NC: \u0111\u00e1nh gi\u00e1 hi\u1ec7u qu\u1ea3 can thi\u1ec7p ph\u1ed5 qu\u00e1t d\u1ef1a tr\u00ean CBT t\u1ea1i tr\u01b0\u1eddng \u0111\u1ec3 gi\u1ea3m lo \u00e2u \u1edf HS l\u1edbp 9.')

p = doc.add_paragraph()
r = p.add_run('Ph\u01b0\u01a1ng ph\u00e1p: Cluster RCT \u2014 36 tr\u01b0\u1eddng ph\u00e2n ng\u1eabu nhi\u00ean (18 can thi\u1ec7p + 18 \u0111\u1ed1i ch\u1ee9ng), m\u1ed7i nh\u00e1nh 360 HS l\u1edbp 9. Can thi\u1ec7p: 8 phi\u00ean CBT/tu\u1ea7n (40 ph\u00fat) b\u1edfi GV \u0111\u01b0\u1ee3c \u0111\u00e0o t\u1ea1o + 1 th\u00e1ng t\u1ef1 th\u1ef1c h\u00e0nh + 1 phi\u00ean cha m\u1eb9. \u0110o: SCARED (lo \u00e2u, \u03b1=0,87), DASS-21 (tr\u1ea7m c\u1ea3m), Rosenberg (t\u1ef1 tr\u1ecdng) t\u1ea1i ban \u0111\u1ea7u, sau can thi\u1ec7p, v\u00e0 3 th\u00e1ng. GEE \u0111i\u1ec1u ch\u1ec9nh clustering + g\u00e2y nhi\u1ec5u.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('K\u1ebft qu\u1ea3: T\u1ef7 l\u1ec7 m\u1ea5t m\u1eabu < 1%. Lo \u00e2u GI\u1ea2M \u0111\u00e1ng k\u1ec3 sau 3 th\u00e1ng: \u03b2 = \u22120,096 (KTC 95%: \u22120,186 \u0111\u1ebfn \u22120,005, p = 0,038). T\u1ef1 tr\u1ecdng T\u0102NG \u0111\u00e1ng k\u1ec3 SAU can thi\u1ec7p: \u03b2 = 0,811 (KTC: 0,314\u20131,309, p = 0,001). Tr\u1ea7m c\u1ea3m kh\u00f4ng c\u00f3 \u00fd ngh\u0129a th\u1ed1ng k\u00ea (OR = 0,257, p = 0,098 sau can thi\u1ec7p; OR = 0,422, p = 0,052 follow-up). Lo \u00e2u kh\u00f4ng gi\u1ea3m ngay sau can thi\u1ec7p (p = 0,115) nh\u01b0ng gi\u1ea3m \u0111\u00e1ng k\u1ec3 sau 3 th\u00e1ng \u2014 g\u1ee3i \u00fd c\u1ea7n TH\u1edcI GIAN \u0111\u1ec3 k\u1ef9 n\u0103ng CBT ph\u00e1t huy t\u00e1c d\u1ee5ng.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

# DANH GIA NHANH
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'Child Adol Psychiatry Mental Health Q1 IF \u2248 4,0. Open Access.',
    'CLUSTER RCT \u2014 thi\u1ebft k\u1ebf M\u1ea0NH NH\u1ea4T cho can thi\u1ec7p tr\u01b0\u1eddng h\u1ecdc. 36 tr\u01b0\u1eddng, 720 HS.',
    'CBT do GV cung c\u1ea5p (kh\u00f4ng ph\u1ea3i chuy\u00ean gia t\u00e2m l\u00fd) \u2014 KH\u1ea2 THI cho LMIC nh\u01b0 VN (thi\u1ebfu chuy\u00ean gia).',
    'HS l\u1edbp 9 (~14 tu\u1ed5i) \u2014 \u0111\u00fang nh\u00f3m tu\u1ed5i c\u1ea5p 2 c\u1ee7a \u0111\u1ec1 t\u00e0i.',
    'Lo \u00e2u GI\u1ea2M sau 3 th\u00e1ng (\u03b2=\u22120,096, p=0,038) \u2014 c\u1ea7n th\u1eddi gian th\u1ef1c h\u00e0nh.',
    'T\u1ef1 tr\u1ecdng T\u0102NG sau can thi\u1ec7p (\u03b2=0,811, p=0,001) \u2014 ph\u00f9 h\u1ee3p Ireland QT32 (t\u1ef1 tr\u1ecdng quan tr\u1ecdng h\u01a1n).',
    'Bao g\u1ed3m PHI\u00caN CHA M\u1eb8 \u2014 ph\u00f9 h\u1ee3p v\u0103n h\u00f3a ch\u00e2u \u00c1 (cha m\u1eb9 tham gia). Islam QT31: cha m\u1eb9 AOR=0,75.',
    'M\u1ea5t m\u1eabu < 1% \u2014 c\u1ef1c th\u1ea5p, ch\u1ea5t l\u01b0\u1ee3ng cao.',
    'Sri Lanka \u2014 LMIC ch\u00e2u \u00c1, b\u1ed1i c\u1ea3nh g\u1ea7n VN h\u01a1n ph\u01b0\u01a1ng T\u00e2y.',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf:', bold=True)
for b in [
    'Ch\u1ec9 Sri Lanka \u2014 kh\u00e1c VN v\u1ec1 gi\u00e1o d\u1ee5c, v\u0103n h\u00f3a.',
    'Lo \u00e2u kh\u00f4ng gi\u1ea3m NGAY sau can thi\u1ec7p (p=0,115) \u2014 ch\u1ec9 sau 3 th\u00e1ng.',
    'Tr\u1ea7m c\u1ea3m kh\u00f4ng c\u1ea3i thi\u1ec7n \u0111\u00e1ng k\u1ec3 (p=0,052\u20130,098).',
    'GV cung c\u1ea5p \u2014 ch\u1ea5t l\u01b0\u1ee3ng kh\u00f4ng b\u1eb1ng chuy\u00ean gia. Kh\u00f4ng m\u00f9 (HS + GV bi\u1ebft nh\u00f3m).',
    'SCARED 41 m\u1ee5c \u2014 d\u00e0i, c\u00f3 th\u1ec3 g\u00e2y m\u1ec7t m\u1ecfi cho HS.',
]:
    add_p(doc, f'\u2022 {b}')

# PHUONG PHAP
add_page_ref(doc, '2\u20136', 'Child Adol Psychiatry MH', '2024, 18:108')
add_heading(doc, '1. PH\u01af\u01a0NG PH\u00c1P', 2)
add_p(doc, '1.1. Thi\u1ebft k\u1ebf: Cluster RCT \u2014 36 tr\u01b0\u1eddng t\u1ea1i qu\u1eadn Colombo, Sri Lanka. Ph\u00e2n t\u1ea7ng theo lo\u1ea1i tr\u01b0\u1eddng (nam/n\u1eef/h\u1ed7n h\u1ee3p). Ph\u00e2n ng\u1eabu nhi\u00ean kh\u1ed1i 1:1: 18 can thi\u1ec7p + 18 \u0111\u1ed1i ch\u1ee9ng. M\u1ed7i tr\u01b0\u1eddng: 1 l\u1edbp 9 ch\u1ecdn ng\u1eabu nhi\u00ean (20 HS). T\u1ed5ng: 360 + 360 = 720 HS. \u0110\u0103ng k\u00fd th\u1eed nghi\u1ec7m: SLCTR/2018/018.')
add_p(doc, '1.2. Can thi\u1ec7p:', bold=True)
add_heading(doc, 'B\u1ea3ng 1. N\u1ed9i dung 8 phi\u00ean CBT t\u1ea1i tr\u01b0\u1eddng', 3)
add_table(doc,
    ['Phi\u00ean', 'N\u1ed9i dung'],
    [['1 (Tu\u1ea7n 1)', 'Gi\u00e1o d\u1ee5c t\u00e2m l\u00fd: nh\u1eadn bi\u1ebft c\u1ea3m x\u00fac, lo \u00e2u l\u00e0 g\u00ec?'],
     ['2 (Tu\u1ea7n 2)', 'Tr\u00e1nh suy ngh\u0129 ti\u00eau c\u1ef1c + th\u1ef1c h\u00e0nh th\u01b0 gi\u00e3n'],
     ['3 (Tu\u1ea7n 3)', 'Tr\u00e1nh suy ngh\u0129 ti\u00eau c\u1ef1c v\u1edbi b\u00e0i t\u1eadp t\u00ecnh hu\u1ed1ng'],
     ['4 (Tu\u1ea7n 4)', 'K\u1ef9 n\u0103ng gi\u1ea3i quy\u1ebft v\u1ea5n \u0111\u1ec1'],
     ['5 (Tu\u1ea7n 5)', 'Gi\u1ea3i quy\u1ebft v\u1ea5n \u0111\u1ec1 n\u00e2ng cao + b\u00e0i t\u1eadp'],
     ['6 (Tu\u1ea7n 6)', 'Ph\u01a1i nhi\u1ec5m (exposure) + qu\u1ea3n l\u00fd lo \u00e2u t\u00ecnh hu\u1ed1ng'],
     ['7 (Tu\u1ea7n 7)', 'K\u1ef9 n\u0103ng x\u00e3 h\u1ed9i + giao ti\u1ebfp quy\u1ebft \u0111o\u00e1n'],
     ['8 (Tu\u1ea7n 8)', '\u00d4n t\u1eadp + ph\u00f2ng ng\u1eeba t\u00e1i ph\u00e1t + k\u1ebf ho\u1ea1ch th\u1ef1c h\u00e0nh'],
     ['Th\u00e1ng 9\u201312', 'T\u1ef1 th\u1ef1c h\u00e0nh d\u01b0\u1edbi gi\u00e1m s\u00e1t GV'],
     ['Phi\u00ean cha m\u1eb9', 'Gi\u00e1o d\u1ee5c v\u1ec1 lo \u00e2u + c\u00e1ch h\u1ed7 tr\u1ee3 con']],
    widths=[3.0, 9.0])
add_p(doc, 'GV \u0111\u01b0\u1ee3c \u0111\u00e0o t\u1ea1o CBT + c\u1ea9m nang + gi\u00e1m s\u00e1t b\u1edfi chuy\u00ean gia SKTT. S\u1ed5 tay GV, v\u1edf b\u00e0i t\u1eadp HS, c\u00f4ng c\u1ee5 gi\u00e1m s\u00e1t th\u1ef1c h\u00e0nh, t\u1edd r\u01a1i cha m\u1eb9.', size=9, italic=True)

add_p(doc, '1.3. C\u00f4ng c\u1ee5 \u0111o:', bold=True)
add_p(doc, '\u2022 SCARED (Screen for Child Anxiety Related Disorders, 41 m\u1ee5c, Likert 3 \u0111i\u1ec3m, \u03b1=0,87). \u0110i\u1ec3m cao = lo \u00e2u l\u1edbn h\u01a1n.')
add_p(doc, '\u2022 DASS-21 ti\u1ec3u thang tr\u1ea7m c\u1ea3m. Ng\u01b0\u1ee1ng: 19 (sensitivity=80%, specificity=83%). \u0110\u00e3 x\u00e1c th\u1ef1c Sinhala.')
add_p(doc, '\u2022 Rosenberg Self-Esteem Scale (10 m\u1ee5c, Likert 4 \u0111i\u1ec3m, 10\u201340). \u0110i\u1ec3m cao = t\u1ef1 tr\u1ecdng cao.')
add_p(doc, '1.4. Ph\u00e2n t\u00edch: ITT. GEE (Generalized Estimating Equation) \u0111i\u1ec1u ch\u1ec9nh clustering tr\u01b0\u1eddng + g\u00e2y nhi\u1ec5u. \u0110o 3 th\u1eddi \u0111i\u1ec3m: ban \u0111\u1ea7u, sau can thi\u1ec7p, 3 th\u00e1ng.')

# KET QUA
add_page_ref(doc, '6\u201310', 'Child Adol Psychiatry MH', '2024, 18:108')
add_heading(doc, '2. K\u1ebeT QU\u1ea2', 2)

add_heading(doc, 'B\u1ea3ng 2. K\u1ebft qu\u1ea3 ch\u00ednh \u2014 GEE (can thi\u1ec7p vs \u0111\u1ed1i ch\u1ee9ng)', 3)
add_table(doc,
    ['K\u1ebft qu\u1ea3', 'Sau can thi\u1ec7p', 'p', 'Sau 3 th\u00e1ng', 'p'],
    [['Lo \u00e2u (SCARED) \u03b2', '\u22120,024 (\u22120,055 \u0111\u1ebfn 0,006)', '0,115', '\u22120,096 (\u22120,186 \u0111\u1ebfn \u22120,005)', '0,038*'],
     ['Tr\u1ea7m c\u1ea3m (DASS-21) OR', '0,257 (0,052\u20131,286)', '0,098', '0,422 (0,177\u20131,008)', '0,052'],
     ['T\u1ef1 tr\u1ecdng (Rosenberg) \u03b2', '0,811 (0,314\u20131,309)', '0,001***', '0,435 (\u22120,276\u20131,145)', '0,231']],
    widths=[3.5, 4.0, 1.5, 4.0, 1.5])
add_p(doc, 'LO \u00c2U: kh\u00f4ng gi\u1ea3m NGAY (p=0,115) nh\u01b0ng GI\u1ea2M \u0111\u00e1ng k\u1ec3 sau 3 th\u00e1ng (p=0,038) \u2014 c\u1ea7n th\u1eddi gian th\u1ef1c h\u00e0nh k\u1ef9 n\u0103ng. T\u1ef0 TR\u1eccNG: T\u0102NG ngay (p=0,001) nh\u01b0ng kh\u00f4ng duy tr\u00ec (p=0,231) \u2014 c\u1ea7n duy tr\u00ec li\u00ean t\u1ee5c. TR\u1ea6M C\u1ea2M: bi\u00ean gi\u1edbi (p=0,052) \u2014 g\u1ea7n \u0111\u00e1ng k\u1ec3.', size=9, italic=True)

# THAO LUAN + KET LUAN
add_heading(doc, '3. TH\u1ea2O LU\u1eacN V\u00c0 K\u1ebeT LU\u1eacN', 2)
add_p(doc, 'CBT ph\u1ed5 qu\u00e1t t\u1ea1i tr\u01b0\u1eddng do GV cung c\u1ea5p HI\u1ec6U QU\u1ea2 gi\u1ea3m lo \u00e2u v\u00e0 t\u0103ng t\u1ef1 tr\u1ecdng \u1edf VTN. \u0110\u1eb7c bi\u1ec7t, lo \u00e2u ch\u1ec9 gi\u1ea3m sau 3 TH\u00c1NG (kh\u00f4ng ph\u1ea3i ngay) \u2014 g\u1ee3i \u00fd k\u1ef9 n\u0103ng CBT c\u1ea7n TH\u1edcI GIAN TH\u1ef0C H\u00c0NH \u0111\u1ec3 ph\u00e1t huy t\u00e1c d\u1ee5ng. Ph\u00f9 h\u1ee3p v\u1edbi BMC NMA QT29 (CBT SUCRA 0,66 \u2014 hi\u1ec7u qu\u1ea3 h\u1ea1ng 2).')
add_p(doc, 'GV cung c\u1ea5p (kh\u00f4ng ph\u1ea3i chuy\u00ean gia) \u2014 m\u00f4 h\u00ecnh KH\u1ea2 THI cho LMIC nh\u01b0 VN n\u01a1i thi\u1ebfu chuy\u00ean gia t\u00e2m l\u00fd. Ph\u00f9 h\u1ee3p UNICEF VN22 (c\u1ea7n \u0111\u00e0o t\u1ea1o GV SKTT). Phi\u00ean cha m\u1eb9 \u2014 ph\u00f9 h\u1ee3p v\u0103n h\u00f3a ch\u00e2u \u00c1 (Islam QT31: cha m\u1eb9 AOR=0,75 b\u1ea3o v\u1ec7). Ireland QT32: OGA (One Good Adult) b\u1ea3o v\u1ec7.')
add_p(doc, 'T\u1ef1 tr\u1ecdng t\u0103ng ngay nh\u01b0ng kh\u00f4ng duy tr\u00ec \u2014 c\u1ea7n can thi\u1ec7p LI\u00caN T\u1ee4C, kh\u00f4ng ch\u1ec9 8 phi\u00ean. Ph\u00f9 h\u1ee3p Ireland QT32 (t\u1ef1 tr\u1ecdng quan tr\u1ecdng h\u01a1n theo th\u1eddi gian).')

# TLTK
add_heading(doc, 'T\u00c0I LI\u1ec6U THAM KH\u1ea2O', 2)
for ref in [
    'De Silva, S. et al. (2024). Effectiveness of CBT-based intervention for reducing anxiety among adolescents in Colombo District, Sri Lanka. Child Adol Psych MH, 18:108.',
    'De Silva, S. et al. (2023). Development and process evaluation of a universal school-based CBT intervention. [Protocol paper]',
    '(Xem \u0111\u1ea7y \u0111\u1ee7 30+ TLTK trong b\u00e0i g\u1ed1c)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('CBT', 'Cognitive Behavioral Therapy \u2014 Li\u1ec7u ph\u00e1p Nh\u1eadn th\u1ee9c\u2013H\u00e0nh vi'),
    ('Cluster RCT', 'Th\u1eed nghi\u1ec7m Ng\u1eabu nhi\u00ean C\u1ee5m c\u00f3 \u0110\u1ed1i ch\u1ee9ng'),
    ('SCARED', 'Screen for Child Anxiety Related Disorders (41 m\u1ee5c)'),
    ('DASS-21', 'Depression Anxiety Stress Scale 21'),
    ('GEE', 'Generalized Estimating Equation'),
    ('ITT', 'Intention-to-Treat'),
    ('LMIC', 'Low and Middle-Income Countries'),
    ('\u03b2', 'Beta \u2014 H\u1ec7 s\u1ed1 h\u1ed3i quy'),
    ('OR', 'Odds Ratio \u2014 T\u1ef7 s\u1ed1 ch\u00eanh'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N / CRITICAL REVIEW')
add_red(doc, '\u0110i\u1ec3m m\u1ea1nh:', bold=True)
for s in [
    'Cluster RCT \u2014 thi\u1ebft k\u1ebf M\u1ea0NH NH\u1ea4T cho can thi\u1ec7p tr\u01b0\u1eddng. 36 tr\u01b0\u1eddng, 720 HS. \u0110\u0103ng k\u00fd th\u1eed nghi\u1ec7m.',
    'GV cung c\u1ea5p CBT (kh\u00f4ng chuy\u00ean gia) \u2014 m\u00f4 h\u00ecnh KH\u1ea2 THI cho LMIC. VN c\u00f3 th\u1ec3 \u00e1p d\u1ee5ng ngay.',
    'HS l\u1edbp 9 (~14 tu\u1ed5i) \u2014 \u0111\u00fang \u0111\u1ed1i t\u01b0\u1ee3ng c\u1ea5p 2 c\u1ee7a \u0111\u1ec1 t\u00e0i. So: VN21 (l\u1edbp 6\u21929, ESSA), QT29 (NMA 30 RCT nhưng ph\u01b0\u01a1ng T\u00e2y).',
    'Phi\u00ean cha m\u1eb9 \u2014 ph\u00f9 h\u1ee3p v\u0103n h\u00f3a VN. Islam QT31: cha m\u1eb9 AOR=0,75.',
    'M\u1ea5t m\u1eabu < 1% \u2014 c\u1ef1c th\u1ea5p. GEE \u0111i\u1ec1u ch\u1ec9nh clustering.',
    'Lo \u00e2u gi\u1ea3m sau 3 th\u00e1ng (kh\u00f4ng ngay) \u2014 nh\u1eadn th\u1ee9c quan tr\u1ecdng cho thi\u1ebft k\u1ebf can thi\u1ec7p.',
    'Open Access CC BY. Q1 IF \u2248 4,0.',
]:
    add_red(doc, f'\u2022 {s}')
add_red(doc, 'H\u1ea1n ch\u1ebf chi ti\u1ebft:', bold=True)
for s in [
    'Ch\u1ec9 Sri Lanka (Colombo) \u2014 kh\u00e1c VN. Ng\u00f4n ng\u1eef Sinhala. Gi\u00e1o d\u1ee5c Anh qu\u1ed1c.',
    'Lo \u00e2u kh\u00f4ng gi\u1ea3m NGAY (p=0,115) \u2014 hi\u1ec7u qu\u1ea3 ch\u1eadm. C\u00f3 th\u1ec3 do 8 phi\u00ean ch\u01b0a \u0111\u1ee7.',
    'Tr\u1ea7m c\u1ea3m bi\u00ean gi\u1edbi (p=0,052) \u2014 g\u1ea7n \u0111\u00e1ng k\u1ec3 nh\u01b0ng kh\u00f4ng \u0111\u1ea1t. C\u1ea7n m\u1eabu l\u1edbn h\u01a1n?',
    'KH\u00d4NG m\u00f9 (HS + GV bi\u1ebft nh\u00f3m) \u2014 thi\u00ean l\u1ec7ch k\u1ef3 v\u1ecdng. Kh\u00f3 m\u00f9 trong can thi\u1ec7p tr\u01b0\u1eddng.',
    'SCARED 41 m\u1ee5c \u2014 d\u00e0i. GAD-7 (7 m\u1ee5c) ng\u1eafn h\u01a1n, d\u1ec5 tri\u1ec3n khai.',
    'Can thi\u1ec7p PHỔ QU\u00c1T (t\u1ea5t c\u1ea3 HS, kh\u00f4ng ch\u1ec9 HS lo \u00e2u) \u2014 hi\u1ec7u qu\u1ea3 th\u1ea5p h\u01a1n can thi\u1ec7p ch\u1ec9 \u0111\u1ecbnh (indicated).',
]:
    add_red(doc, f'\u2022 {s}')
add_red(doc, 'Kho\u1ea3ng tr\u1ed1ng NC / Gap:', bold=True)
for s in [
    'VN C\u1ea6N cluster RCT CBT tr\u01b0\u1eddng t\u01b0\u01a1ng t\u1ef1 \u2014 \u0111\u00e2y l\u00e0 M\u1eaau T\u1ed0T NH\u1ea4T cho \u0111\u1ec1 c\u01b0\u01a1ng giai \u0111o\u1ea1n 2. M\u00f4 h\u00ecnh: 8 phi\u00ean CBT/tu\u1ea7n + 1 th\u00e1ng th\u1ef1c h\u00e0nh + phi\u00ean cha m\u1eb9.',
    'Th\u00edch \u1ee9ng VN: d\u1ecbch ti\u1ebfng Vi\u1ec7t + v\u00ed d\u1ee5 t\u00ecnh hu\u1ed1ng VN (thi v\u00e0o l\u1edbp 10, \u00e1p l\u1ef1c gia \u0111\u00ecnh) + t\u00edch h\u1ee3p gi\u00e1 tr\u1ecb VN (ph\u00f9 h\u1ee3p B7 CA-CBT \u0110NA).',
    'K\u1ebft h\u1ee3p CBT + PE (ho\u1ea1t \u0111\u1ed9ng th\u1ec3 ch\u1ea5t) \u2014 BMC NMA QT29: PE SUCRA=0,51 c\u0169ng hi\u1ec7u qu\u1ea3.',
    'So s\u00e1nh: can thi\u1ec7p ph\u1ed5 qu\u00e1t (t\u1ea5t c\u1ea3 HS) vs ch\u1ec9 \u0111\u1ecbnh (ch\u1ec9 HS lo \u00e2u) t\u1ea1i VN.',
    'Theo d\u00f5i d\u00e0i h\u01a1n: 6 th\u00e1ng, 12 th\u00e1ng \u2014 lo \u00e2u c\u00f3 DUY TR\u00cc gi\u1ea3m kh\u00f4ng?',
]:
    add_red(doc, f'\u2022 {s}')

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '42_SriLanka_CBT_RCT_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['36 tr\u01b0\u1eddng', '720', '0,038', '0,811', '0,001', 'SCARED', 'DASS-21', 'Colombo', 'cluster']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
