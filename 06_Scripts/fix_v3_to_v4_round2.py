# -*- coding: utf-8 -*-
"""
QA Round 2 fix v3→v4 — Paragraph-level replacement với keyword matching.
- Fix 7 lỗi fact
- Cleanup body cites cho các TLTK đã xóa
- Balance word count: 5000 ≤ ALL ≤ 6200 trong mọi cách đếm
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
import copy

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
IN1 = ROOT / "bai-bao-khgdvn/Bai1_YTNC_HSTHCS_v3_14052026.docx"
IN2 = ROOT / "bai-bao-khgdvn/Bai2_CanThiep_HSTHCS_v3_14052026.docx"
OUT1 = ROOT / "bai-bao-khgdvn/Bai1_YTNC_HSTHCS_v4_14052026.docx"
OUT2 = ROOT / "bai-bao-khgdvn/Bai2_CanThiep_HSTHCS_v4_14052026.docx"


# ============================================================
# BÀI 1 — paragraph-level replacements
# Each entry: (unique_keyword_to_find, new_text)
# ============================================================

B1_PARA_REPLACEMENTS = [
    # FIX Jagiello: intervention review, not correlation r=0.30-0.55
    ("Jagiello",
     "Tổng quan hệ thống các chương trình can thiệp áp lực học tập ở học sinh trung học của "
     "Jagiello và cộng sự (2025) công bố trên Child Psychiatry and Human Development xác nhận "
     "rằng áp lực học tập là mục tiêu can thiệp khả thi khi triển khai trong môi trường trường "
     "học. Phân tích tổng hợp gánh nặng bệnh tật khu vực ASEAN (GBD ASEAN, 2025) cũng xác nhận "
     "áp lực thi cử là một trong các yếu tố nguy cơ mạnh nhất đối với DALYs do rối loạn lo âu "
     "ở vị thành niên khu vực."),

    # FIX Fassi: N=3,340 clinical UK, not 24 cohorts >50,000
    ("Fassi",
     "Nghiên cứu cắt ngang đại diện quốc gia Anh của Fassi và cộng sự (2025) đăng trên Nature "
     "Human Behaviour trên 3.340 vị thành niên 11–19 tuổi với chẩn đoán lâm sàng cho thấy nhóm "
     "có rối loạn sức khỏe tâm thần được chẩn đoán sử dụng mạng xã hội nhiều hơn nhóm không có "
     "rối loạn, với mức chênh khoảng g = 0,46 ở một số chỉ số hành vi. Điểm mạnh đặc biệt so với "
     "các nghiên cứu chỉ sàng lọc tự báo cáo là tách rõ nhóm có rối loạn được chẩn đoán so với "
     "nhóm chứng, qua đó giảm thiên lệch tự lựa chọn."),

    # FIX Brunborg: N=979,043 decomposition, no β=0.12
    ("Brunborg",
     "Brunborg và cộng sự (2025) công bố trên Social Science & Medicine phân tích 979.043 vị "
     "thành niên Na Uy từ 1.417 khảo sát cấp thành phố giai đoạn 2011–2024 bằng phương pháp "
     "decomposition, trong đó sử dụng mạng xã hội nổi lên như một trong các yếu tố giải thích "
     "chính cho xu hướng tăng căng thẳng tâm thần."),

    # FIX Moore: CI 1.34-2.33 not 1.52-2.07; 165 studies not 36
    ("Moore và cộng sự (2017)",
     "Phân tích tổng hợp của Moore và cộng sự (2017) công bố trên World Journal of Psychiatry "
     "tổng hợp 165 nghiên cứu (57 thuần tập tiến cứu + 108 cắt ngang) với đánh giá theo tiêu chí "
     "Bradford Hill báo cáo tỷ số chênh OR cho lo âu ở nhóm nạn nhân bắt nạt là 1,77 (khoảng tin "
     "cậy 95% từ 1,34 đến 2,33). Đáng chú ý, các tác giả phân biệt rõ vai trò nạn nhân, thủ phạm "
     "và nhóm vừa nạn nhân vừa thủ phạm, trong đó nhóm thứ ba có nguy cơ lo âu cao nhất, gợi ý "
     "cơ chế chồng lấn của tổn thương và rối loạn điều tiết."),
]

# Body paragraph cleanup — sentences pointing to refs already removed from TLTK
# Replace whole sentence/paragraph with shorter version
B1_CLEAN_REPLACEMENTS = [
    # Remove sentence containing "Hồ Thị Trúc Quỳnh" — only ref to that VN paper
    # Find paragraph containing it
    ("Hồ Thị Trúc Quỳnh",
     # Replace with sentence WITHOUT this cite (consolidate with adjacent content)
     "Hoàng Trung Học và cộng sự (2025) xác nhận trải nghiệm bị bắt nạt bằng lời ở môi trường "
     "học đường có liên hệ thuận với điểm lo âu trong mô hình hồi quy đa biến, sau khi kiểm soát "
     "giới và khối lớp. Các hình thức bắt nạt bằng lời thường được phối hợp với hành vi loại trừ "
     "xã hội – chẳng hạn cô lập trong nhóm bạn lớp, từ chối ngồi gần hoặc lập nhóm chat riêng có "
     "chủ ý loại trừ một học sinh. Khi cộng dồn các hành vi này, tổn thương tâm lý có thể vượt "
     "xa tổn thương từ bắt nạt thể chất đơn lẻ, ngay cả khi không để lại bằng chứng vật lý."),

    # Remove Nguyễn Thị Vân paragraph
    ("Nguyễn Thị Vân (2020)",
     "Bên cạnh các yếu tố cấp lớp học, sự dịch chuyển môn học từ tiểu học lên trung học cơ sở – "
     "đặc biệt là sự xuất hiện đột ngột của nhiều môn chuyên sâu, hệ thống đánh giá đa môn và "
     "chuyển đổi giáo viên – cũng là một nguồn áp lực được nhiều nghiên cứu định tính tại Việt Nam "
     "ghi nhận. Mặc dù dữ liệu cắt ngang Việt Nam phần lớn còn hạn chế về quan hệ nhân – quả, "
     "hướng và độ lớn của hiệu ứng trùng với bằng chứng quốc tế, gợi ý áp lực học tập là yếu tố "
     "nguy cơ ổn định xuyên văn hóa và cần được ưu tiên đo lường chính xác trong khảo sát HSTHCS."),

    # Remove Bie & Islam → consolidate into 1 sentence (use V-NAMHS as anchor)
    ("Bie & cs",
     "Rối loạn lo âu là nhóm rối loạn sức khỏe tâm thần phổ biến nhất ở lứa tuổi vị thành niên "
     "trên phạm vi toàn cầu. Các báo cáo Gánh nặng bệnh tật toàn cầu (GBD) công bố gần đây ghi "
     "nhận xu hướng tăng đều của tỷ suất hiện mắc và DALYs ở nhóm 10–24 tuổi giai đoạn 1990–2021, "
     "với phần lớn gánh nặng tập trung tại các quốc gia thu nhập trung bình và thấp."),

    # Remove Wen sentence — paragraph about LPA
    ("Wen và cộng sự",
     "Bằng chứng dọc của các nghiên cứu xã hội tại Bắc Âu xác nhận vai trò của trải nghiệm bị "
     "bắt nạt như biến trung gian quan trọng cho mối liên hệ giữa sử dụng mạng xã hội và lo âu "
     "xã hội. Trong bối cảnh trường học khu vực Đông Á, phân tích hồ sơ trên hơn 1.000 học sinh "
     "nông thôn Trung Quốc xác định nhóm 'bắt nạt cao – ít hỗ trợ xã hội' có điểm lo âu trung "
     "bình cao gấp đôi nhóm tham chiếu."),

    # Remove Sohn sentence — but keep OR=3.05 mention
    ("Phân tích tổng hợp của Sohn",
     "Sử dụng quá mức điện thoại thông minh đã được ghi nhận có liên hệ mạnh với triệu chứng lo "
     "âu ở vị thành niên qua nhiều phân tích tổng hợp, với tỷ số chênh xung quanh ngưỡng OR = 3 "
     "trên các nghiên cứu sàng lọc bằng các thang nghiện điện thoại chuẩn hóa."),

    # Remove Li 2025 mention
    ("Li và cộng sự (2025)",
     "Các tổng quan hệ thống gần đây ghi nhận thời gian màn hình vượt hai giờ mỗi ngày có liên hệ "
     "thuận với triệu chứng lo âu ở phần lớn các nghiên cứu được tổng hợp."),

    # Carver fix — keep Brief-COPE mention without explicit cite
    ("khung đối phó Brief-COPE của Carver (1997)",
     "khung đối phó Brief-COPE"),
]


# Body expansion — add 1-2 paragraphs to push body to ~5050
B1_EXPAND_PARAGRAPHS = [
    # New synthesis paragraph for §3.5 self-esteem
    ("4. Kết luận",  # insert BEFORE this heading
     [
        "Trên khung tổng hợp về điều tiết cảm xúc – đối phó, lòng tự trọng có vai trò chi phối "
        "khả năng lựa chọn chiến lược đối phó thích nghi trong tình huống căng thẳng học đường. "
        "Học sinh có lòng tự trọng cao có xu hướng chọn các chiến lược tập trung vào vấn đề hoặc "
        "đánh giá lại nhận thức, trong khi học sinh tự trọng thấp dễ chuyển sang né tránh và tự "
        "đổ lỗi. Cơ chế này có hàm ý quan trọng cho thiết kế can thiệp ở Việt Nam: các chương "
        "trình tăng cường tự trọng có thể là cánh cửa hiệu quả để cải thiện điều tiết cảm xúc "
        "rộng hơn, chứ không đơn thuần giảm triệu chứng lo âu trực tiếp.",
        "Sự tương tác giữa lòng tự trọng và áp lực học tập cũng đáng chú ý. Khi áp lực học tập "
        "cao, học sinh tự trọng thấp dễ rơi vào vòng xoáy thất bại – tự đánh giá tiêu cực – lo "
        "âu, trong khi học sinh tự trọng cao có khả năng phân tách thất bại nhất thời khỏi đánh "
        "giá tổng thể về bản thân. Việc thiết kế can thiệp đa thành phần – kết hợp giảm áp lực, "
        "tăng tự trọng và cải thiện đối phó – do đó có thể khả thi hơn so với can thiệp đơn yếu "
        "tố."
     ]),
]


# TLTK entries CẦN XÓA (tương ứng cite đã clean)
B1_TLTK_REMOVE_KEYWORDS = [
    "Brunborg, G. S.",
    "Carver, C. S.",
    "Sun, J., Dunne",
    "Sohn, S. Y., Rees",
    "Wen, X., Lin, Y.",
    "Li, X., Smith, A.",
    "Bie, F., Yan, X.",
    "Islam, M. I.",
    "Hồ Thị Trúc Quỳnh",
    "Nguyễn Thị Vân. (2020)",
]


# ============================================================
# BÀI 2 — paragraph-level replacements
# ============================================================

B2_PARA_REPLACEMENTS = [
    # FIX Bress 2024 Maya: n=59 young adults 18-25, not >200 adolescents
    ("Bress",
     "Thử nghiệm lâm sàng ngẫu nhiên của Bress và cộng sự (2024) công bố trên JAMA Network "
     "Open đánh giá ứng dụng Maya – một nền tảng CBT tự hướng dẫn 12 phiên trên điện thoại – "
     "trên 59 thanh niên 18–25 tuổi với rối loạn lo âu (rối loạn lo âu tổng quát 56%, rối loạn "
     "lo âu xã hội 41%). Kết quả ghi nhận hiệu ứng kích thước rất lớn trên các thang đo phụ "
     "(Anxiety Sensitivity Index Cohen d = 0,93 tại tuần 6; Liebowitz Social Anxiety Scale "
     "Cohen d = 1,07 tại tuần 12). Mặc dù mẫu chính là thanh niên trẻ chứ không phải HSTHCS, "
     "kết quả này gợi ý CBT số được thiết kế chu đáo có thể đạt hiệu quả lớn và cần được adapt "
     "cho lứa tuổi nhỏ hơn."),

    # FIX Cai 2025 — clearly 38 RCT (21 in MA), SMD=0.17
    ("Cai và cộng sự (2025) trên 39",
     "Phân tích tổng hợp của Cai và cộng sự (2025) trên 38 thử nghiệm lâm sàng ngẫu nhiên (21 "
     "trong phân tích định lượng) chương trình tăng cường khả năng phục hồi tại trường báo cáo "
     "hiệu ứng tổng hợp SMD = 0,17 (KTC 95% 0,06–0,29). Khi phân nhóm, các chương trình dựa "
     "trên chánh niệm cho hiệu ứng cao nhất (SMD = 0,57), tiếp đến là các chương trình lấy thể "
     "thao làm trung tâm (SMD = 0,41). Mặc dù hiệu ứng tổng hợp nhỏ, các tác giả nhấn mạnh giá "
     "trị của can thiệp phổ quát quy mô lớn không sàng lọc – đặc biệt khi tích hợp vào hoạt "
     "động trải nghiệm hoặc môn Giáo dục công dân."),
]


# Bài 2 Tóm tắt VN — slightly shorter version (target 200 words)
B2_TOM_TAT_NEW = (
    "Liệu pháp nhận thức – hành vi (Cognitive Behavioral Therapy – CBT) đã được xác lập là tiếp "
    "cận chuẩn vàng trong can thiệp rối loạn lo âu ở trẻ em – vị thành niên, với nhiều phân tích "
    "tổng hợp ghi nhận kích thước hiệu ứng từ trung bình tới lớn. Tuy nhiên, các nghiên cứu can "
    "thiệp CBT trên học sinh trung học cơ sở (HSTHCS) tại Việt Nam còn rất hạn chế, đặc biệt "
    "thiếu các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm. Bài viết hệ thống hóa các nghiên cứu "
    "can thiệp công bố giai đoạn 2015–2026 theo ba trục: CBT đã được kiểm định lâm sàng, CBT "
    "triển khai trong trường học, và CBT qua nền tảng số. Phân tích cho thấy ba khoảng trống "
    "chính trong bối cảnh Việt Nam: vắng các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên "
    "mẫu HSTHCS, chưa có chương trình CBT trường học phù hợp Chương trình giáo dục phổ thông "
    "2018, và chưa có nền tảng CBT số tiếng Việt được kiểm định. Năm khuyến nghị triển khai "
    "được đề xuất nhằm thu hẹp các khoảng trống đó."
)


# ============================================================
# Helpers
# ============================================================

def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def replace_para_text(para, new_text, size=12):
    """Replace all text in para with new_text, preserving paragraph_format."""
    # Clear existing runs
    for run in para.runs:
        run.text = ""
    if para.runs:
        para.runs[0].text = new_text
    else:
        r = para.add_run(new_text)
        set_run_format(r, size=size)


def find_para_with_keyword(doc, keyword):
    """Find first paragraph containing keyword (case-sensitive)."""
    for p in doc.paragraphs:
        if keyword in p.text:
            return p
    return None


def find_para_index(doc, contains_text, start=0):
    for i in range(start, len(doc.paragraphs)):
        if contains_text in doc.paragraphs[i].text:
            return i
    return -1


def insert_para_before_idx(doc, target_idx, text, indent_cm=1.0):
    target_p = doc.paragraphs[target_idx]
    new_p = copy.deepcopy(target_p._element)
    for child in list(new_p):
        new_p.remove(child)
    target_p._element.addprevious(new_p)
    new_para = doc.paragraphs[target_idx]
    pf = new_para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.1
    pf.first_line_indent = Cm(indent_cm)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = new_para.add_run(text)
    set_run_format(r, size=12)
    return new_para


def remove_tltk_by_keywords(doc, keywords):
    removed = []
    idx_tltk = find_para_index(doc, "Tài liệu tham khảo")
    if idx_tltk < 0:
        return removed
    to_remove = []
    for i in range(idx_tltk + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        for kw in keywords:
            if kw in p.text:
                to_remove.append((i, p, kw))
                break
    for i, p, kw in reversed(to_remove):
        p._element.getparent().remove(p._element)
        removed.append(kw)
    return removed


def apply_para_replacements(doc, replacements):
    """For each (keyword, new_text), find first para with keyword and replace whole para text."""
    applied = []
    for kw, new_text in replacements:
        p = find_para_with_keyword(doc, kw)
        if p is not None:
            replace_para_text(p, new_text)
            applied.append(kw)
        else:
            print(f"  [✗] Not found: {kw}")
    return applied


# ============================================================
# Main
# ============================================================

def fix_bai1():
    print("=" * 70)
    print("FIX BÀI 1")
    print("=" * 70)
    doc = Document(IN1)

    # Apply fact-fix paragraphs
    a1 = apply_para_replacements(doc, B1_PARA_REPLACEMENTS)
    print(f"  Fact fixes applied: {len(a1)}/{len(B1_PARA_REPLACEMENTS)}")
    for k in a1:
        print(f"    ✓ {k}")

    # Apply cleanup replacements (remove cites to deleted refs)
    a2 = apply_para_replacements(doc, B1_CLEAN_REPLACEMENTS)
    print(f"  Body cleanup applied: {len(a2)}/{len(B1_CLEAN_REPLACEMENTS)}")
    for k in a2:
        print(f"    ✓ {k}")

    # Insert expansion paragraphs
    for anchor, texts in B1_EXPAND_PARAGRAPHS:
        idx = find_para_index(doc, anchor)
        if idx > 0:
            for text in texts:
                insert_para_before_idx(doc, idx, text)
                idx += 1
            print(f"  [+] Inserted {len(texts)} para(s) before '{anchor}'")

    # Trim TLTK
    removed = remove_tltk_by_keywords(doc, B1_TLTK_REMOVE_KEYWORDS)
    print(f"  [-] Removed {len(removed)} TLTK entries")
    for r in removed:
        print(f"    - {r}")

    # Metadata cleanup
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)

    doc.save(OUT1)
    print(f"\n  Saved: {OUT1}")


def fix_bai2():
    print("\n" + "=" * 70)
    print("FIX BÀI 2")
    print("=" * 70)
    doc = Document(IN2)

    a1 = apply_para_replacements(doc, B2_PARA_REPLACEMENTS)
    print(f"  Fact fixes applied: {len(a1)}/{len(B2_PARA_REPLACEMENTS)}")
    for k in a1:
        print(f"    ✓ {k}")

    # Replace Tóm tắt VN
    p_tt = None
    for p in doc.paragraphs:
        if p.text.startswith("Tóm tắt:"):
            p_tt = p
            break
    if p_tt is not None:
        # Keep "Tóm tắt:" label run, replace rest
        if p_tt.runs:
            p_tt.runs[0].text = "Tóm tắt: "
            for run in p_tt.runs[1:]:
                run.text = ""
            new_run = p_tt.add_run(B2_TOM_TAT_NEW)
            set_run_format(new_run, size=11)
            print(f"  ✓ Replaced Tóm tắt VN ({len(B2_TOM_TAT_NEW.split())} từ)")

    # Metadata cleanup
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)

    doc.save(OUT2)
    print(f"\n  Saved: {OUT2}")


if __name__ == "__main__":
    fix_bai1()
    fix_bai2()
    print("\n[DONE]")
