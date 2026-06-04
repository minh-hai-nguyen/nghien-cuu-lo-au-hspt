"""DOC: Tra loi cau hoi cua thay:
'Neu Gia thuyet Yeu to nguy co tu nghien dien thoai duoc kiem chung thi
nghien cuu se lam gi?'

Format: CTH v6 + CAU TRA LOI to xanh, dap ung pattern cac gia thuyet khac.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/GT_nghien_dien_thoai_duoc_kiem_chung_thi_lam_gi.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('NẾU GIẢ THUYẾT NGHIỆN ĐIỆN THOẠI\nLÀ YẾU TỐ NGUY CƠ ĐƯỢC KIỂM CHỨNG\nTHÌ NGHIÊN CỨU SẼ LÀM GÌ?\n— Trả lời câu hỏi của thầy 09/05/2026 —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi của thầy
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Nếu Giả thuyết Yếu tố nguy cơ đến từ NGHIỆN ĐIỆN THOẠI được '
    'kiểm chứng thì nghiên cứu sẽ làm gì? Các giả thuyết khác đều '
    'trả lời câu này rồi.', italic=True
)

# Phát biểu lại giả thuyết
H('1. Phát biểu lại giả thuyết H — Nghiện điện thoại', level=2, color=NAVY)
para(
    'Học sinh trung học cơ sở có MỨC ĐỘ NGHIỆN ĐIỆN THOẠI cao hơn '
    'sẽ có TỶ LỆ rối loạn lo âu cao hơn — sau khi kiểm soát các '
    'biến nhân khẩu (giới, lớp, kinh tế gia đình) và các yếu tố '
    'tâm lý xã hội đồng biến (giấc ngủ, áp lực học tập, quan hệ '
    'cha mẹ-con).', bold=True
)
para(
    'Nghiện điện thoại (problematic smartphone use) đo bằng SAS-SV '
    '(Smartphone Addiction Scale-Short Version, Kwon và cộng sự '
    '2013) hoặc BSMAS (Bergen Social Media Addiction Scale, '
    'Andreassen và cộng sự 2016).'
)
para(
    'Trong chương 3 luận án CTH, kết quả β chuẩn hóa của NĐT '
    'với ba dạng RLLA (Bảng 3.26): NĐT → RLLATQ = 0,336; NĐT → '
    'RLLAXH = 0,383 (mạnh nhất); NĐT → RLLACL = 0,265. β tổng '
    '= 0,400; R² = 0,160. Tất cả p < 0,001.'
)

# 2. Phù hợp với pattern các giả thuyết khác
H('2. Mô hình trả lời chuẩn — bảy hành động sau khi giả thuyết được kiểm chứng', level=2, color=NAVY)
para(
    'Mô phỏng cấu trúc đã dùng cho các giả thuyết khác (ALHT, tự '
    'trọng, bắt nạt, ...). Nếu giả thuyết NĐT được kiểm chứng — '
    'tức β NĐT → RLLA có ý nghĩa thống kê và cường độ đáng kể — '
    'nghiên cứu sẽ tiến hành BẢY HÀNH ĐỘNG sau:'
)

H('Hành động 1 — Phát triển công cụ sàng lọc NĐT chuẩn hóa cho HS THCS Việt Nam', level=3)
para(
    'Hiện tại Việt Nam chưa có thang nghiện điện thoại được chuẩn '
    'hóa cho học sinh trung học cơ sở. Chương 3 luận án dùng '
    'thang riêng. Hành động: dịch + chuẩn hóa SAS-SV (Kwon 2013) '
    'hoặc BSMAS (Andreassen 2016) cho mẫu HS THCS VN, kiểm tra '
    'độ tin cậy (Cronbach α ≥ 0,80) và độ giá trị (CFA fit '
    'indices CFI ≥ 0,95). Đây là ĐÓNG GÓP PHƯƠNG PHÁP LUẬN '
    'quan trọng cho nghiên cứu VN.'
)

H('Hành động 2 — Đề xuất ngưỡng cắt cho sàng lọc trường học', level=3)
para(
    'Sau khi chuẩn hóa thang đo, đề xuất NGƯỠNG CẮT cho phân '
    'loại "có dấu hiệu nghiện" vs "bình thường" — dựa trên '
    'ROC analysis với outcome RLLA. Áp dụng ngưỡng để sàng lọc '
    'học sinh có nguy cơ cao trong các đợt khám sức khỏe học '
    'đường định kỳ. Tham chiếu: Sohn và cộng sự (2019) trong '
    'BMC Psychiatry phân tích tổng hợp 41 nghiên cứu N = 41.871 '
    'cho thấy tỷ lệ nghiện điện thoại trung vị 23,3% (KTC 95% '
    '14,0%–31,2%) — gợi ý ngưỡng phân loại 25% học sinh trong '
    'mẫu VN.'
)

H('Hành động 3 — Thiết kế can thiệp giảm screen time tại trường', level=3)
para(
    'Phối hợp với chương trình Khung tập huấn (mục 3.8): bổ sung '
    'CHỦ ĐỀ 4 "Kiểm soát sử dụng điện thoại" thành CAN THIỆP '
    'TRỌNG TÂM cho nhóm có nguy cơ cao. Áp dụng cấu trúc:'
)
para('• Theo dõi screen time bằng app điện thoại (Module 2 Clear Fear).')
para('• Cam kết "giờ không điện thoại" (ví dụ 1 giờ trước khi ngủ).')
para('• Đào tạo kỹ năng tự kiểm soát: kỹ thuật Pomodoro, "đặt điện thoại trong tủ".')
para('• Nội dung digital alternative: hoạt động ngoài trời, đọc sách, thể thao.')
para(
    'Phù hợp với Schmidt-Persson và cộng sự (2024) trong JAMA '
    'Network Open: thử nghiệm RCT giảm screen media use trong 2 '
    'tuần cho 89 gia đình Đan Mạch giảm khó khăn nội hóa với '
    'Cohen d = 0,53 (effect size trung bình).'
)

H('Hành động 4 — Tích hợp can thiệp digital (iCBT app) cho lo âu xã hội', level=3)
para(
    'Vì β NĐT → RLLAXH = 0,383 là MẠNH NHẤT trong ba dạng RLLA '
    'đối với nghiện điện thoại, đề xuất can thiệp DIGITAL ĐẶC '
    'THÙ cho lo âu xã hội. Phù hợp với phân tích tổng hợp 21 '
    'RCT của Walder và cộng sự (2025) trong Journal of Medical '
    'Internet Research:'
)
para('• DMHI tổng quát: g = 0,508 (medium)')
para('• DMHI ĐẶC THÙ cho SAD: g = 0,878 (large) — gần gấp đôi DMHI tổng')
para('• DMHI có HƯỚNG DẪN người: g = 0,825 vs DMHI không hướng dẫn g ≈ 0,2')
para(
    'Đề xuất: phát triển app iCBT TIẾNG VIỆT cho lo âu xã hội '
    'với thành phần GUIDED (chat-bot hoặc tin nhắn từ giáo viên/'
    'cố vấn). Nội dung dựa trên 6 module Clear Fear (Samele và '
    'cộng sự, 2025): psychoeducation, symptom tracking, '
    'cognitive challenges, relaxation, behavioral activation, '
    'exposure ladder.'
)

H('Hành động 5 — Đào tạo cha mẹ về quản lý thời gian điện tử', level=3)
para(
    'Phù hợp với mô hình EACP của Lochman và cộng sự (2025): '
    '16 buổi 90 phút cho cha mẹ với nội dung cốt lõi:'
)
para('• Hiểu cơ chế nghiện điện thoại + so sánh xã hội + FOMO ở thanh thiếu niên.')
para('• Kỹ năng đặt giới hạn không phán xét: "không điện thoại trong bữa ăn", "trả điện thoại trước 9h tối".')
para('• Mô hình hóa hành vi: cha mẹ giảm screen time của chính mình.')
para('• Đối thoại không phán xét về nội dung con xem trên mạng xã hội.')
para('• Xác định hoạt động thay thế: thể thao, đọc sách, hoạt động gia đình.')
para(
    'Phù hợp với phát hiện chương 3: β quan hệ cha mẹ-con tổng (HTCM) = '
    '−0,234 cho RLLA — yếu tố bảo vệ TRUNG BÌNH–CAO. Tăng cường '
    'hỗ trợ cha mẹ vừa giảm trực tiếp lo âu vừa giảm gián tiếp '
    'qua giảm thời gian điện tử.'
)

H('Hành động 6 — Phân tích trung gian (mediation) tìm cơ chế cụ thể', level=3)
para(
    'Sau khi xác lập β NĐT → RLLA có ý nghĩa, đề xuất phân tích '
    'TRUNG GIAN để tìm CƠ CHẾ:'
)
para('• Mô hình 1: NĐT → giấc ngủ kém → RLLA (cơ chế trục HPA)')
para('• Mô hình 2: NĐT → giảm self-efficacy → RLLA (theo Zheng & Peng 2025: 63,13% qua self-efficacy)')
para('• Mô hình 3: NĐT → so sánh xã hội tiêu cực → RLLAXH (theo Festinger 1954 + Fassi 2025)')
para('• Mô hình 4: NĐT → giảm tương tác mặt đối mặt → RLLAXH')
para(
    'Phân tích này cung cấp HƯỚNG can thiệp cụ thể: nếu cơ chế '
    'chính là self-efficacy, can thiệp tập trung vào tăng tự '
    'tin học sinh; nếu cơ chế chính là giấc ngủ, can thiệp tập '
    'trung vào "giờ không điện thoại trước khi ngủ".'
)

H('Hành động 7 — Đối chiếu với y văn quốc tế và đề xuất chính sách', level=3)
para(
    'Đối chiếu cường độ β chương 3 (NĐT → RLLA tổng = 0,400) '
    'với phát hiện quốc tế:'
)
add_table(
    ['Nguồn', 'Cỡ mẫu', 'Đo', 'Kết quả'],
    [
        ['Sohn 2019 BMC Psychiatry meta', 'N = 41.871', 'PSU', 'OR lo âu = 3,05 (KTC 2,64–3,53; I² = 0%)'],
        ['Chen 2023 (QT007)', 'n = 63.205 TQ', 'IGD ≥ 32', 'OR = 5,00 (mạnh thứ 2)'],
        ['Zheng & Peng 2025 (QT041)', 'HS nghề TQ', 'SMAS', 'r = 0,415; β = 0,153; mediation 63,13%'],
        ['Hoàng Trung Học 2025 (VN014)', 'n = 8.473 VN', 'Thời gian dùng', 'β = 0,176'],
        ['Li 2025 (QT022) BJCP', 'n = 4.058 Úc dọc', 'Screen time', 'Cắt ngang sig; DỌC NS (p = 0,443)'],
        ['Chương 3 luận án CTH', 'n = 1.352 VN THCS', 'Thang riêng', 'β = 0,400; R² = 0,160'],
    ]
)
para('')
para(
    'Đề xuất CHÍNH SÁCH cấp Bộ Giáo dục:'
)
para('• Hướng dẫn quốc gia về sàng lọc nghiện điện thoại định kỳ tại trường THCS.')
para('• Tích hợp thành phần "kỹ năng số lành mạnh" vào chương trình giáo dục công dân lớp 6–9.')
para('• Đào tạo giáo viên tư vấn học đường về phát hiện sớm và can thiệp với học sinh có dấu hiệu nghiện điện thoại.')
para('• Phối hợp với Bộ Y tế phát triển app iCBT tiếng Việt — có thể tham khảo Clear Fear (Samele 2025) làm khung.')

# 3. Cảnh báo về causal direction
H('3. Cảnh báo: bằng chứng dọc cho thấy quan hệ HAI CHIỀU', level=2, color=NAVY)
para(
    'Một bằng chứng QUAN TRỌNG đối lại với giả thuyết một chiều: '
    'Li và cộng sự (2025) — QT022 — trong British Journal of '
    'Clinical Psychology với thiết kế DỌC 12 tháng trên 4.058 '
    'thanh thiếu niên Úc phát hiện screen time T1 KHÔNG dự báo '
    'lo âu T2 (p = 0,443). Trái với mối liên quan cắt ngang đáng '
    'kể, mối liên quan dọc YẾU/KHÔNG đáng kể.'
)
para(
    'Nói cách khác, có khả năng quan hệ HAI CHIỀU: vị thành niên '
    'có triệu chứng lo âu/trầm cảm tăng dùng điện thoại (rút lui '
    'xã hội, phân tâm) — không chỉ ngược lại. Đây là khả năng '
    'REVERSE CAUSATION cần cảnh báo trong giả thuyết.'
)
para(
    'Hệ quả cho hành động: KHÔNG chỉ tập trung "giảm điện thoại" '
    '— phải kết hợp can thiệp song hành GIẢM lo âu (qua các '
    'thành phần CBT đã đề xuất) để tránh hiệu ứng vòng xoáy.'
)

# 4. CÂU TRẢ LỜI tô xanh
H('4. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue_run('Nếu giả thuyết NĐT là yếu tố nguy cơ được kiểm chứng, nghiên cứu sẽ làm BẢY việc:', bold=True)
blue_run(
    '(1) PHÁT TRIỂN CÔNG CỤ SÀNG LỌC NĐT chuẩn hóa cho HS THCS '
    'VN — dịch + chuẩn hóa SAS-SV (Kwon 2013) hoặc BSMAS '
    '(Andreassen 2016); kiểm tra Cronbach α + CFA fit indices.'
)
blue_run(
    '(2) ĐỀ XUẤT NGƯỠNG CẮT cho sàng lọc trường học — dựa trên '
    'ROC analysis với outcome RLLA; tham chiếu Sohn 2019 (tỷ '
    'lệ trung vị 23,3%) → ngưỡng ~25% học sinh có nguy cơ.'
)
blue_run(
    '(3) THIẾT KẾ CAN THIỆP GIẢM SCREEN TIME tại trường — '
    'theo dõi screen time + cam kết "giờ không điện thoại" + '
    'kỹ năng tự kiểm soát + hoạt động thay thế. Phù hợp '
    'Schmidt-Persson 2024 RCT giảm screen time → cải thiện '
    'SKTT (Cohen d = 0,53).'
)
blue_run(
    '(4) TÍCH HỢP iCBT APP cho LO ÂU XÃ HỘI — vì β NĐT → '
    'RLLAXH = 0,383 mạnh nhất. Phù hợp Walder 2025: DMHI '
    'SAD-specific g = 0,878 (gần gấp đôi tổng quát); guided '
    'g = 0,825 vs unguided g ≈ 0,2. Dùng 6 module Clear Fear '
    '(Samele 2025) làm khung.'
)
blue_run(
    '(5) ĐÀO TẠO CHA MẸ về quản lý thời gian điện tử — 16 '
    'buổi 90 phút theo mô hình EACP (Lochman 2025); kỹ năng '
    'đặt giới hạn không phán xét + mô hình hóa hành vi + '
    'đối thoại về nội dung mạng xã hội.'
)
blue_run(
    '(6) PHÂN TÍCH TRUNG GIAN tìm cơ chế — 4 mô hình: '
    'NĐT→giấc ngủ→RLLA; NĐT→self-efficacy→RLLA (Zheng 2025: '
    '63,13%); NĐT→so sánh xã hội→RLLAXH; NĐT→giảm tương '
    'tác trực tiếp→RLLAXH. Cơ chế chính sẽ định hướng can '
    'thiệp cụ thể.'
)
blue_run(
    '(7) ĐỐI CHIẾU Y VĂN QUỐC TẾ + ĐỀ XUẤT CHÍNH SÁCH — '
    'so sánh với Sohn 2019 (OR=3,05), Chen 2023 (IGD '
    'OR=5,00), Zheng 2025 (mediation 63,13%). Đề xuất '
    'sàng lọc quốc gia + tích hợp "kỹ năng số lành mạnh" '
    'vào chương trình lớp 6–9 + phát triển app iCBT tiếng '
    'Việt.'
)
blue_run('Cảnh báo bằng chứng đối lại:', bold=True)
blue_run(
    'Li 2025 (QT022) BJCP n=4.058 Úc dọc 12 tháng — screen '
    'time T1 KHÔNG dự báo lo âu T2 (p = 0,443). Có khả năng '
    'quan hệ HAI CHIỀU (lo âu → tăng dùng điện thoại). Phải '
    'kết hợp can thiệp giảm điện thoại VỚI giảm lo âu — '
    'tránh hiệu ứng vòng xoáy. Phát biểu CHÍNH XÁC: "MỐI '
    'LIÊN QUAN" thay vì "GÂY RA".'
)

# 5. TLTK
H('5. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Andreassen, C. S., Billieux, J., Griffiths, M. D., Kuss, D. J., Demetrovics, Z., Mazzoni, E., & Pallesen, S. (2016). The relationship between addictive use of social media and video games and symptoms of psychiatric disorders. Psychology of Addictive Behaviors, 30(2), 252–262. https://doi.org/10.1037/adb0000160',
    'Chen, Z., Ren, S., He, R., và cộng sự. (2023). Prevalence and associated factors of depressive and anxiety symptoms among Chinese secondary school students. BMC Psychiatry, 23(1), 580. [QT007.]',
    'Festinger, L. (1954). A theory of social comparison processes. Human Relations, 7(2), 117–140.',
    'Fassi, L., và cộng sự. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9(6), 1283–1299. [QT027.]',
    'Hoàng, T. H., & Nguyễn, T. D. (2025). Mức độ căng thẳng, lo âu và trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 tại Việt Nam. [VN014.]',
    'Kwon, M., Kim, D. J., Cho, H., & Yang, S. (2013). The Smartphone Addiction Scale: Development and validation of a short version for adolescents. PLoS ONE, 8(12), e83558. https://doi.org/10.1371/journal.pone.0083558',
    'Li, S. H., và cộng sự. (2025). Cross-sectional and longitudinal associations of screen time with adolescent depression and anxiety. British Journal of Clinical Psychology, 64(4), 873–887. [QT022.]',
    'Lochman, J. E., và cộng sự. (2025). Randomized controlled trial of the Early Adolescent Coping Power program. Journal of School Psychology. [QT065.]',
    'Samele, C., và cộng sự. (2025). Clear Fear app for adolescents. JMIR Formative Research. [QT062.]',
    'Schmidt-Persson, J., và cộng sự. (2024). Screen media use and mental health of children and adolescents: A secondary analysis of a randomized clinical trial. JAMA Network Open, 7(7), e2419881. [QT033.]',
    'Sohn, S. Y., Rees, P., Wildridge, B., Kalk, N. J., & Carter, B. (2019). Prevalence of problematic smartphone usage and associated mental health outcomes amongst children and young people: A systematic review, meta-analysis and GRADE of the evidence. BMC Psychiatry, 19(1), 356. https://doi.org/10.1186/s12888-019-2350-x',
    'Walder, N., Frey, A., Berger, T., và cộng sự. (2025). Digital mental health interventions for prevention and treatment of social anxiety disorder in children, adolescents and young adults. Journal of Medical Internet Research. [QT040.]',
    'Zheng, G. F., & Peng, H. Y. (2025). The effects of social media addiction, academic stress, and sleep quality on anxiety symptoms. Psychology Research and Behavior Management, 18, 1571–1584. [QT041.]',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
