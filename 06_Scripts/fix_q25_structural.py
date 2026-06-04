# -*- coding: utf-8 -*-
"""
Xử lý 8 lỗi cấu trúc bài Q2.5 (EN v4 + VN v4 -> v5). Giữ v4 để đối chứng.
1. Biến "school attachment / gắn bó trường học" chưa khai trong Phương pháp -> khai + cite Goodenow (1993).
2. So sánh β với OR (EN para 24) -> sửa, nêu rõ hai thước đo không cùng hệ quy chiếu.
3. Phương pháp trùng lặp -> bỏ khai báo AI lặp lại trong đoạn "Phân tích thống kê".
4. Tên phần mềm SEM mơ hồ -> chèn placeholder tô đỏ để tác giả ghi rõ.
5. Bảng 12 tạp chí trong file EN còn tiếng Việt -> dịch sang tiếng Anh.
6. Abstract gọi nhầm yếu tố bảo vệ là "risk factor" -> tách rõ nguy cơ / bảo vệ.
7. Ngôn ngữ nhân-quả -> làm mềm động từ nhân quả ("generates"/"tạo ra").
8. Bản VN thiếu đoạn cảnh báo SEM cắt ngang -> chèn đoạn mới.
"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor

RED = RGBColor(0xFF, 0x00, 0x00)


def rep(doc, idx, old, new, label):
    """Thay chuỗi trong 1 run của đoạn idx — giữ định dạng."""
    p = doc.paragraphs[idx]
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            print(f"  [{idx}] OK: {label}")
            return True
    # thử ghép run
    full = p.text
    if old in full and p.runs:
        p.runs[0].text = full.replace(old, new)
        for r in p.runs[1:]:
            r.text = ""
        print(f"  [{idx}] OK (gộp run): {label}")
        return True
    print(f"  [{idx}] !! KHÔNG khớp: {label} -> {old[:50]}")
    return False


def split_redmark(doc, idx, phrase, placeholder, drop=None):
    """Dựng lại đoạn idx: bỏ chuỗi drop (nếu có), thay phrase bằng placeholder tô đỏ."""
    p = doc.paragraphs[idx]
    txt = p.text
    if drop:
        txt = txt.replace(drop, "")
    if phrase not in txt:
        print(f"  [{idx}] !! KHÔNG khớp phrase phần mềm")
        return False
    before, after = txt.split(phrase, 1)
    name = p.runs[0].font.name if p.runs else "Times New Roman"
    size = p.runs[0].font.size if p.runs else Pt(12)
    for r in list(p.runs):
        r.text = ""
    # run 1
    r1 = p.runs[0]
    r1.text = before
    r1.font.name = name; r1.font.size = size
    r2 = p.add_run(placeholder)
    r2.font.name = name; r2.font.size = size; r2.font.color.rgb = RED
    r3 = p.add_run(after)
    r3.font.name = name; r3.font.size = size
    print(f"  [{idx}] OK: tô đỏ placeholder phần mềm SEM")
    return True


def set_cell(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


# ================= EN =================
print("=== Q25 EN v4 -> v5 ===")
en = Document("bai-bao-khgdvn/Q25_SEM_Pathways_EN_v4.docx")

# 6 + 1: abstract — phân loại nguy cơ/bảo vệ + thêm school attachment
rep(en, 1,
    "linking five school-relevant risk factors (academic pressure, smartphone addiction, school bullying, self-esteem, parental support) and one coping factor to three anxiety subtypes (generalised, social, separation)",
    "linking school-based risk factors (academic pressure, smartphone addiction, school bullying), protective factors (self-esteem, perceived parental support, school attachment) and one coping factor to three anxiety subtypes (generalised, social, separation)",
    "abstract: risk/protective + school attachment")
rep(en, 1,
    "perceived parental support scale, and a brief coping inventory.",
    "perceived parental support scale, a school-membership scale, and a brief coping inventory.",
    "abstract: measures list + school-membership scale")

# 1 + 6: Q1 trong Đặt vấn đề
rep(en, 8,
    "Do the five school-relevant risk factors (academic pressure, smartphone addiction, school bullying, self-esteem, parental support) show differential pathways to the three anxiety subtypes?",
    "Do the school-relevant risk and protective factors (academic pressure, smartphone addiction, school bullying, self-esteem, parental support, school attachment) show differential pathways to the three anxiety subtypes?",
    "Q1: risk+protective + school attachment")

# 1: khai school attachment trong Measures
rep(en, 15,
    "perceived parental support with the parental subscale of the Multidimensional Scale of Perceived Social Support, and coping strategies with a brief inventory inspired by Carver's (1997) Brief COPE.",
    "perceived parental support with the parental subscale of the Multidimensional Scale of Perceived Social Support, school attachment with the Psychological Sense of School Membership scale (Goodenow, 1993), and coping strategies with a brief inventory inspired by Carver's (1997) Brief COPE.",
    "Measures: khai school attachment + Goodenow")

# 1: khai school attachment trong Detailed measurement
rep(en, 20,
    "measured self-esteem and parental support respectively. The Brief COPE was used",
    "measured self-esteem and parental support respectively. School attachment was measured with the Psychological Sense of School Membership scale (Goodenow, 1993). The Brief COPE was used",
    "Detailed measurement: school attachment")

# 2: β vs OR
rep(en, 24,
    "The β = 0.376 pathway from school bullying to separation anxiety, by contrast, lies well above the odds ratio of 1.77 documented by Moore and colleagues for general anxiety outcomes, reinforcing the subtype-specific reading of our finding.",
    "The β = 0.376 pathway from school bullying to separation anxiety cannot be placed on the same metric as the odds ratio of 1.77 reported by Moore and colleagues for general anxiety outcomes, because standardised regression coefficients and odds ratios are not directly comparable; the present subtype-specific finding nonetheless refines that meta-analytic evidence by localising the bullying–anxiety association within the separation-anxiety subtype.",
    "β-vs-OR comparison fixed")

# 7: làm mềm ngôn ngữ nhân quả
rep(en, 44,
    "in which repeated victimisation in the school environment generates persistent avoidance of separation from primary caregivers as a safe-haven strategy",
    "in which repeated victimisation in the school environment is associated with persistent avoidance of separation from primary caregivers as a safe-haven strategy",
    "causal language softened (generates->associated with)")

# 3 + 4: bỏ khai báo AI lặp + tô đỏ tên phần mềm
AI_EN = (" Throughout, we adopt the convention recommended by the Vietnam Journal of "
         "Educational Sciences for transparent declaration of artificial-intelligence "
         "assistance: the authors used a large-language-model tool only for "
         "English-language polishing and reference formatting; all analytic decisions, "
         "interpretation, and writing of substantive content were carried out by the "
         "human authors.")
split_redmark(en, 18, "widely used SEM software",
              "[SEM software to be specified by the authors: AMOS, Mplus, or lavaan]",
              drop=AI_EN)

# 5: dịch Bảng 2 (12 tạp chí) sang tiếng Anh
tb = en.tables[2]
EN_TBL = {
 (1,2):"3.10",(1,4):"~3,150 CHF (LMIC waiver)",(1,5):"~77 days",(1,7):"Top submission priority — speed and prestige",
 (2,2):"2.77",(2,4):"Editor's discretion (LMIC waiver)",(2,5):"Not reported",(2,7):"Strongest cultural fit — suited to the Asian region",
 (3,2):"2.20",(3,4):"~$2–3k",(3,5):"Not reported",(3,7):"Springer + PubMed indexed; accepts SEM studies",
 (4,2):"3.40",(4,4):"€1,390 (OA)",(4,5):"Not reported",(4,7):"Cheapest OA among the Q1–Q2 options; PMC indexed",
 (5,2):"2.14",(5,4):"Subscription (free for authors)",(5,5):"~8.7 weeks",(5,7):"Good fit for policy implications and school services",
 (6,2):"3.27",(6,4):"~$2,500",(6,5):"~77 days",(6,7):"Suitable when emphasising developmental psychology",
 (7,2):"~2.5",(7,4):"Hybrid (optional OA)",(7,5):"Not reported",(7,7):"Taylor & Francis hybrid model — review the policy carefully before submitting",
 (8,2):"2.04",(8,5):"Not reported",(8,7):"Emphasises school-based implications",
 (9,2):"~3.0",(9,5):"Not reported",(9,7):"Exact niche for school-based mental health",
 (10,2):"2.60",(10,4):"$1,695 (OA)",(10,5):"~185 days",(10,7):"High acceptance rate ~50%; broad scope",
 (11,2):"~3.5",(11,4):"$2,500 (OA)",(11,5):"Not reported",(11,7):"LMIC focus — Vietnamese authors welcome",
 (12,2):"~1.3",(12,5):"Not reported",(12,7):"Backup option if the cascade is rejected from higher tiers"}
ncell = 0
for (r, c), v in EN_TBL.items():
    set_cell(tb.rows[r].cells[c], v)
    ncell += 1
print(f"  Bảng 2: dịch {ncell} ô sang tiếng Anh")

cp = en.core_properties
cp.author = ""; cp.last_modified_by = ""
en.save("bai-bao-khgdvn/Q25_SEM_Pathways_EN_v5.docx")
print("  -> đã lưu Q25_SEM_Pathways_EN_v5.docx\n")

# ================= VN =================
print("=== Q25 VN v4 -> v5 ===")
vn = Document("bai-bao-khgdvn/Q25_SEM_Pathways_VN_v4.docx")

# 6 + 1: abstract
rep(vn, 1,
    "gắn năm yếu tố nguy cơ học đường (áp lực học tập, nghiện điện thoại thông minh, bắt nạt học đường, lòng tự trọng, hỗ trợ cha mẹ) và một yếu tố đối phó với ba phân nhóm lo âu (tổng quát, xã hội, chia ly)",
    "gắn các yếu tố nguy cơ học đường (áp lực học tập, nghiện điện thoại thông minh, bắt nạt học đường), các yếu tố bảo vệ (lòng tự trọng, hỗ trợ cha mẹ tri giác, gắn bó trường học) và một yếu tố đối phó với ba phân nhóm lo âu (tổng quát, xã hội, chia ly)",
    "abstract: tách nguy cơ/bảo vệ + gắn bó trường học")
rep(vn, 1,
    "thang hỗ trợ cha mẹ tri giác và thang đối phó ngắn.",
    "thang hỗ trợ cha mẹ tri giác, thang gắn bó trường học và thang đối phó ngắn.",
    "abstract: danh mục công cụ + thang gắn bó trường học")

# 1 + 6: Q1
rep(vn, 8,
    "Năm yếu tố nguy cơ học đường (áp lực học tập, nghiện điện thoại thông minh, bắt nạt học đường, lòng tự trọng, hỗ trợ cha mẹ) có cho đường dẫn khác biệt đến ba phân nhóm lo âu không?",
    "Các yếu tố nguy cơ và bảo vệ học đường (áp lực học tập, nghiện điện thoại thông minh, bắt nạt học đường, lòng tự trọng, hỗ trợ cha mẹ, gắn bó trường học) có cho đường dẫn khác biệt đến ba phân nhóm lo âu không?",
    "Q1: nguy cơ+bảo vệ + gắn bó trường học")

# 1: khai gắn bó trường học trong Công cụ đo
rep(vn, 14,
    "và đối phó bằng bảng kiểm ngắn phỏng theo Brief COPE của Carver (1997).",
    "gắn bó trường học bằng thang Cảm nhận về sự gắn kết với trường học (Goodenow, 1993), và đối phó bằng bảng kiểm ngắn phỏng theo Brief COPE của Carver (1997).",
    "Công cụ đo: khai gắn bó trường học + Goodenow")

# 1: khai gắn bó trường học trong Mô tả công cụ chi tiết
rep(vn, 19,
    "đo lường lòng tự trọng và hỗ trợ cha mẹ tương ứng. Brief COPE phiên bản ngắn 16 mục",
    "đo lường lòng tự trọng và hỗ trợ cha mẹ tương ứng. Gắn bó trường học được đo bằng thang Cảm nhận về sự gắn kết với trường học (Goodenow, 1993). Brief COPE phiên bản ngắn 16 mục",
    "Mô tả công cụ: gắn bó trường học")

# 7: làm mềm ngôn ngữ nhân quả
rep(vn, 37,
    "trong đó trải nghiệm bị bắt nạt lặp lại trong môi trường trường học tạo ra hành vi né tránh kéo dài khỏi việc tách rời người chăm sóc chính",
    "trong đó trải nghiệm bị bắt nạt lặp lại trong môi trường trường học đi kèm với hành vi né tránh kéo dài việc tách rời người chăm sóc chính",
    "làm mềm ngôn ngữ nhân quả (tạo ra->đi kèm)")

# 3 + 4: bỏ khai báo AI lặp + tô đỏ tên phần mềm
AI_VN = (" Theo khuyến nghị minh bạch của Tạp chí Khoa học Giáo dục Việt Nam về khai báo "
         "sử dụng trí tuệ nhân tạo, nhóm tác giả có dùng công cụ mô hình ngôn ngữ lớn chỉ "
         "ở vai trò hỗ trợ ngôn ngữ tiếng Anh và định dạng trích dẫn; toàn bộ quyết định "
         "phân tích, diễn giải và viết nội dung khoa học do nhóm tác giả tự thực hiện.")
split_redmark(vn, 17, "phần mềm SEM thông dụng",
              "[phần mềm SEM — tác giả ghi rõ: AMOS, Mplus hoặc lavaan]",
              drop=AI_VN)

# 8: chèn đoạn cảnh báo SEM cắt ngang trước "4. Bàn luận"
# tìm chỉ số đoạn "4. Bàn luận"
disc_idx = None
for i, p in enumerate(vn.paragraphs):
    if p.text.strip() == "4. Bàn luận":
        disc_idx = i
        break
caveat = ("Lưu ý phương pháp luận về SEM cắt ngang. Các phân tích trong nghiên cứu này có "
          "thiết kế cắt ngang, do đó các mô hình cấu trúc tuyến tính được báo cáo ở đây nên "
          "được hiểu là mô tả cấu trúc hiệp phương sai (covariance) giữa các biến số chứ "
          "không phải những khẳng định nhân quả có hướng. Nhóm tác giả theo thông lệ phổ "
          "biến trong tâm lý học phát triển là báo cáo các hệ số đường dẫn chuẩn hóa dạng "
          "hồi quy như một cách tóm tắt cô đọng cho liên hệ cắt ngang, đồng thời nêu rõ sự "
          "mơ hồ về chiều tác động trong phần Bàn luận. Riêng hệ số rất lớn của đối phó kém "
          "thích nghi nên được hiểu là dấu hiệu của hiệp phương sai mạnh với triệu chứng lo "
          "âu chứ không phải một hiệu ứng nhân quả một chiều; mức phù hợp mô hình kém của "
          "đặc tả SEM một chiều củng cố sự thận trọng diễn giải này.")
if disc_idx is not None:
    newp = vn.paragraphs[disc_idx].insert_paragraph_before()
    run = newp.add_run(caveat)
    run.font.name = "Times New Roman"
    run.font.size = Pt(12)
    print(f"  chèn đoạn cảnh báo SEM cắt ngang trước đoạn [{disc_idx}] '4. Bàn luận'")
else:
    print("  !! không tìm thấy '4. Bàn luận'")

cp = vn.core_properties
cp.author = ""; cp.last_modified_by = ""
vn.save("bai-bao-khgdvn/Q25_SEM_Pathways_VN_v5.docx")
print("  -> đã lưu Q25_SEM_Pathways_VN_v5.docx")
print("[DONE]")
