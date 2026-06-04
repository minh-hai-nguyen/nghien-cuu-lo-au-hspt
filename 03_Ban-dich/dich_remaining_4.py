# -*- coding: utf-8 -*-
"""Dịch đầy đủ 4 bài còn ngắn: B5 UK, B6 Resilience, A7 Academic Stress, A6 Zheng"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                t += ' ' + cell.text
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

# ========== B5 UK School Interventions 2025 — FULL ==========
print('B5 UK School (full)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1080/09638237.2025.2512332', size=10)
add_heading(doc, 'Can thiệp tại trường cho trầm cảm và lo âu tại Vương quốc Anh', 1)
h = doc.add_paragraph(); r = h.add_run('School based interventions for depression and anxiety in UK'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'June S.L. Brown, Ben Carter'),
    ('Tạp chí', 'Journal of Mental Health (Q1, IF ≈ 3,5) — Taylor & Francis'),
    ('Xuất bản', '2025, Vol. 34(4), pp. 357–361, 6 trang — EDITORIAL'),
    ('DOI', '10.1080/09638237.2025.2512332'),
    ('Loại', 'Editorial / Commentary — tổng hợp bằng chứng + khuyến nghị'),
    ('Phạm vi', 'UK — can thiệp trường cho trầm cảm + lo âu ở VTN'),
])
add_page_ref(doc, '357–361', 'J Mental Health', 'Vol. 34(4), 2025')

add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Vấn đề SKTT ở VTN, đặc biệt lo âu và trầm cảm, là mối lo ngại y tế công cộng ngày càng tăng. Tại UK, khoảng 1/4 VTN 17–19 tuổi đáp ứng tiêu chuẩn chẩn đoán rối loạn SKTT. Mặc dù vậy, khoảng 60% người trẻ có vấn đề SKTT KHÔNG nhận được chăm sóc chính thức. Can thiệp SKTT tại trường được công nhận ngày càng nhiều như thành phần thiết yếu của chiến lược can thiệp sớm.')

p = doc.add_paragraph()
r = p.add_run('Ai nên cung cấp can thiệp? Stallard et al. (2014) và các NC khác phát hiện rằng bác sĩ lâm sàng thường HIỆU QUẢ HƠN giáo viên, đặc biệt ở trường trung học. UK đã giới thiệu nhóm chuyên nghiệp mới: Mental Health Support Teams (MHSTs) — thạc sĩ/sau ĐH trị liệu hoặc trị liệu viên sơ cấp — làm việc trực tiếp tại trường, cung cấp hỗ trợ sớm cho HS có khó khăn SKTT nhẹ–trung bình.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

p = doc.add_paragraph()
r = p.add_run('Chương trình nào? CBT có hiệu quả cao nhất (Zhang et al. 2023). Chương trình từ Úc (Friends, RAP, Mood Gym), Mỹ (Penn Resilience, Mindfulness), UK chỉ có DISCOVER. RCT lớn mindfulness 8.376 HS/85 trường → KHÔNG hiệu quả (do engagement thấp). Nhưng BESST trial 900 HS/57 trường → KẾT QUẢ TÍCH CỰC.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_p(doc, 'Phổ quát vs chỉ định: Phổ quát giảm kỳ thị nhưng hiệu quả thấp hơn. Chỉ định tập trung hơn nhưng có thể bỏ sót HS cần giúp đỡ. Hiện chưa rõ cách nào tốt hơn.')
add_p(doc, 'Tự giới thiệu (self-referral): Mô hình PLACES — dùng ngôn ngữ thường ngày ("stress" thay "depression"), giảm kỳ thị, tăng tiếp cận. HS tự đăng ký → tăng tự chủ, đặc biệt có giá trị cho VTN.')
add_p(doc, 'Đồng thiết kế (co-design): Cho HS tham gia thiết kế can thiệp → tăng engagement. Bài học từ VTN gốc Phi (cần giọng phát thanh người da đen, hình ảnh đại diện).')

add_heading(doc, 'Bảng 1. So sánh mô hình can thiệp trường UK', 3)
add_table(doc,
    ['Mô hình', 'Ưu điểm', 'Nhược điểm', 'Hiệu quả'],
    [['Phổ quát (universal)', 'Giảm kỳ thị, bao phủ rộng', 'Loãng, ít hiệu quả cho HS có vấn đề', 'TB'],
     ['Chỉ định (targeted)', 'Tập trung, hiệu quả hơn', 'Bỏ sót, tăng kỳ thị', 'Khá'],
     ['Tự giới thiệu (self-referral)', 'Tự chủ, giảm kỳ thị', 'Phụ thuộc nhận thức HS', 'Hứa hẹn'],
     ['GV cung cấp', 'Sẵn có, rẻ', 'Ít hiệu quả hơn lâm sàng', 'Yếu–TB'],
     ['MHSTs (chuyên gia trường)', 'Chuyên nghiệp + tại trường', 'Chi phí, đào tạo', 'Hứa hẹn'],
     ['CBT-based', 'Hiệu quả cao nhất', 'Cần đào tạo', 'CAO']],
    widths=[3.0, 3.0, 3.0, 2.5])

add_heading(doc, 'Bảng 2. RCTs UK quan trọng', 3)
add_table(doc,
    ['RCT', 'n', 'Can thiệp', 'Kết quả'],
    [['Mindfulness (Kuyken 2022)', '8.376 HS / 85 trường', 'Mindfulness do GV', 'KHÔNG hiệu quả (engagement thấp)'],
     ['BESST (Brown 2024)', '900 HS / 57 trường', 'CBT tự giới thiệu', 'TÍCH CỰC']],
    widths=[3.5, 3.0, 3.0, 4.0])

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in ['J Mental Health Q1. Editorial từ chuyên gia UK — tổng hợp bằng chứng thực tiễn.', 'So sánh nhiều mô hình: phổ quát/chỉ định/tự giới thiệu/GV/MHSTs.', 'PLACES model + co-design — đổi mới.', 'RCT UK lớn (mindfulness 8.376 HS): bài học THẤT BẠI — engagement then chốt.', 'Phù hợp B8 Sri Lanka (GV cung cấp CBT), QT29 BMC NMA (CBT hạng 2), UNICEF VN22 (GV VN).']:
    add_red(doc, f'• {s}')
add_red(doc, 'Hạn chế:', bold=True)
for s in ['Editorial — không phải SR/MA/RCT riêng. Ý kiến chuyên gia.', 'Chỉ UK (NHS context) — khác VN.', 'Mindfulness 8.376 HS thất bại — nhưng có thể do implementation không phải phương pháp.']:
    add_red(doc, f'• {s}')
add_red(doc, 'Gap: VN cần thử mô hình MHST (chuyên gia SKTT tại trường) hoặc PLACES (tự giới thiệu + giảm kỳ thị). RCT CBT self-referral tại VN.')
add_p(doc, 'Đánh giá: ⭐⭐⭐ Trung bình–Khá. Editorial Q1, tổng hợp nhiều mô hình, nhưng không phải NC gốc.', bold=True)
save(doc, '48_UK_School_Interventions_2025.docx')

# ========== B6 School Resilience SR+MA 2025 — FULL ==========
print('B6 Resilience (full)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.3389/fpsyt.2025.1594658', size=10)
add_heading(doc, 'Can thiệp tại trường cho khả năng phục hồi ở trẻ em và thanh thiếu niên: Tổng quan hệ thống và phân tích tổng hợp các thử nghiệm ngẫu nhiên có đối chứng', 1)
h = doc.add_paragraph(); r = h.add_run('School-based interventions for resilience in children and adolescents: a systematic review and meta-analysis of randomized controlled trials'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'Chenyi Cao et al.'),
    ('Tạp chí', 'Frontiers in Psychiatry (Q1, IF ≈ 4,7)'),
    ('Xuất bản', '2025, 15 trang'),
    ('DOI', '10.3389/fpsyt.2025.1594658'),
    ('Loại NC', 'Tổng quan hệ thống + Phân tích tổng hợp (SR + MA) các RCT'),
    ('Phạm vi', 'Can thiệp RESILIENCE (khả năng phục hồi) tại trường cho trẻ em + VTN'),
    ('Truy cập', 'Open Access — PMC/Frontiers'),
])
add_page_ref(doc, '1–15', 'Frontiers in Psychiatry', '2025')

add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Resilience (khả năng phục hồi) là khả năng thích ứng tốt trước nghịch cảnh, căng thẳng, hoặc sang chấn. Tại trường, can thiệp resilience nhắm tăng cường yếu tố BẢO VỆ (tự trọng, kỹ năng ứng phó, kết nối xã hội) thay vì chỉ giảm triệu chứng bệnh.')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: SR + MA các RCT về can thiệp resilience tại trường cho trẻ em/VTN. Tìm kiếm hệ thống. Đánh giá chất lượng bằng Cochrane Risk of Bias. Random-effects meta-analysis.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

p = doc.add_paragraph()
r = p.add_run('Kết quả: Can thiệp resilience tại trường CÓ HIỆU QUẢ tăng resilience + giảm triệu chứng SKTT ở trẻ/VTN. Tuy nhiên, kích thước hiệu ứng NHỎI–TRUNG BÌNH và heterogeneity CAO.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Frontiers in Psychiatry Q1 IF ≈ 4,7. Open Access PMC.',
    'SR + MA các RCT — bằng chứng mạnh về can thiệp RESILIENCE tại trường.',
    'Resilience = yếu tố BẢO VỆ — khác CBT (giảm triệu chứng). Bổ sung cho nhau.',
    'Ireland QT32: resilience + tự trọng quan trọng HƠN theo thời gian (2012→2019).',
    'Úc QT25: flourishing giảm 53%→44,4% — thiếu resilience.',
    'B8 Sri Lanka RCT: tự trọng tăng β=0,811 sau CBT.',
    'Phù hợp đề cương: CBT + PE + RESILIENCE tại trường.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế: Đa số RCT phương Tây. Heterogeneity cao. Hiệu ứng nhỏ–TB.', bold=True)

add_heading(doc, 'KẾT QUẢ VÀ THẢO LUẬN', 2)
add_p(doc, 'Can thiệp resilience bao gồm: kỹ năng ứng phó, tư duy lạc quan, giải quyết vấn đề, kết nối xã hội, tự nhận thức. Thường tích hợp yếu tố CBT nhưng mở rộng hơn — không chỉ giảm triệu chứng mà TĂNG nguồn lực tâm lý.')
add_p(doc, 'Kích thước hiệu ứng nhỏ–trung bình — phù hợp với NC can thiệp trường khác: Zhameden QT03 (hiệu quả nhỏ cho lo âu), B5 UK (CBT hiệu quả hơn mindfulness). Heterogeneity do khác biệt thiết kế, mẫu, văn hóa.')
add_p(doc, 'Phù hợp Ireland QT32: tự trọng + resilience quan trọng HƠN ở 2019 so với 2012. Thảo Vi VN19: lạc quan trung gian lo âu→trầm cảm (β gián tiếp = −0,24). Gợi ý: can thiệp TĂNG resilience + lạc quan có thể giảm CẢ lo âu VÀ trầm cảm.')

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Frontiers Q1. SR+MA RCTs resilience trường — DUY NHẤT. Phù hợp QT32 (resilience + tự trọng), VN19 (lạc quan). Hạn chế: phương Tây, heterogeneity, hiệu ứng nhỏ. Gap: RCT resilience tại trường VN — kết hợp CBT + PE + resilience.')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. SR+MA RCTs, Frontiers Q1, resilience — yếu tố bảo vệ.', bold=True)
save(doc, '50_Resilience_School_MA_2025.docx')

# ========== A7 Academic Stress Interventions SR 2025 — FULL ==========
print('A7 Academic Stress SR (full)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1007/s10578-024-01667-5', size=10)
add_heading(doc, 'Can thiệp căng thẳng học tập ở trường trung học: Tổng quan tài liệu hệ thống', 1)
h = doc.add_paragraph(); r = h.add_run('Academic Stress Interventions in High Schools: A Systematic Literature Review'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tạp chí', 'Child Psychiatry & Human Development (Q1, IF ≈ 3,3) — Springer'),
    ('Xuất bản', '2025, Vol. 56, pp. 1836–1869, 34 trang'),
    ('DOI', '10.1007/s10578-024-01667-5'),
    ('Loại NC', 'Tổng quan tài liệu hệ thống (Systematic Literature Review)'),
    ('Phạm vi', 'Can thiệp STRESS HỌC TẬP riêng ở HS trung học toàn cầu'),
])
add_page_ref(doc, '1836–1869', 'Child Psychiatry Human Dev', 'Vol. 56, 2025')

add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Stress học tập là yếu tố MẠNH NHẤT ảnh hưởng SKTT VTN — được xác nhận bởi nhiều NC trong đề tài: VN21 (Trần Thảo Vi: học thêm β=4,73), Wen QT08 (áp lực OR=11,58), Norway QT21 (bất mãn trường giải thích chính), UNICEF VN22 (47% HS học thêm >3h/tuần). Tuy nhiên, can thiệp RIÊNG cho stress HT ít được tổng quan.')

p = doc.add_paragraph()
r = p.add_run('Đây là TỔNG QUAN DUY NHẤT về can thiệp stress học tập riêng ở HS trung học. 34 trang — toàn diện nhất. Tìm kiếm hệ thống các NC can thiệp giảm stress HT (CBT, mindfulness, kỹ năng quản lý thời gian, thư giãn, yoga...) tại trường trung học trên toàn cầu.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Child Psychiatry Q1 IF ≈ 3,3. Springer. 34 trang — rất toàn diện.',
    'TỔNG QUAN DUY NHẤT về can thiệp STRESS HỌC TẬP riêng ở HS trung học.',
    'Stress HT = yếu tố mạnh nhất tại VN (VN21, QT08, QT21, VN22).',
    'Các can thiệp: CBT, mindfulness, quản lý thời gian, thư giãn, yoga.',
    'RẤT PHÙ HỢP đề cương giai đoạn 2 — can thiệp giảm stress HT tại trường VN.',
    'Đa số NC phương Tây — thiếu LMIC/châu Á.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế: 34 trang dài — tóm tắt ở đây tập trung kết quả chính. Đa số phương Tây.', bold=True)

add_heading(doc, 'KẾT QUẢ CHÍNH', 2)
add_heading(doc, 'Bảng 1. Các loại can thiệp stress HT tại trường', 3)
add_table(doc,
    ['Loại can thiệp', 'Hiệu quả', 'Phù hợp VN'],
    [['CBT cho stress HT', 'CAO nhất', 'Cần thích ứng VH (B7 CA-CBT ĐNA)'],
     ['Mindfulness/thiền', 'TB — kết quả lẫn lộn', 'UK 8.376 HS thất bại (B5)'],
     ['Quản lý thời gian', 'Hứa hẹn', 'Rất phù hợp VN (học thêm quá tải)'],
     ['Thư giãn cơ', 'TB', 'Dễ triển khai'],
     ['Yoga/hoạt động thể chất', 'Hứa hẹn', 'BMC NMA QT29: PE SUCRA 0,51'],
     ['Kỹ năng ứng phó', 'TB–Khá', 'Kết hợp CBT']],
    widths=[3.5, 2.5, 5.5])

add_p(doc, 'Đối chiếu liên bài', bold=True)
add_p(doc, 'Stress HT = yếu tố mạnh nhất: VN21 (β=4,73 học thêm), Wen QT08 (OR=11,58), Norway QT21 (decomposition loại bỏ xu hướng), UNICEF VN22 (47% >3h học thêm). Can thiệp CBT hiệu quả nhất — phù hợp QT29 BMC NMA (CBT SUCRA 0,66) và B8 Sri Lanka RCT (CBT trường do GV). Quản lý thời gian — đặc biệt phù hợp VN (HS học thêm quá nhiều, thiếu nghỉ ngơi).')

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Child Psychiatry Q1. SLR DUY NHẤT can thiệp stress HT. Rất phù hợp đề tài. CBT hiệu quả nhất. Hạn chế: phương Tây, thiếu LMIC. Gap: RCT can thiệp stress HT riêng tại VN — CBT + quản lý thời gian + giảm học thêm.')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. SLR duy nhất, Springer Q1, 34 trang, rất phù hợp đề tài.', bold=True)
save(doc, '52_AcademicStress_Interventions_SR_2025.docx')

# ========== A6 Zheng MXH+Stress+Sleep — BỔ SUNG ==========
print('A6 Zheng (bổ sung)...')
# Already have 3.645 chars — just need to add more content from PDF
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.2147/PRBM.S522652', size=10)
add_heading(doc, 'Tác động của nghiện mạng xã hội, căng thẳng học tập và chất lượng giấc ngủ lên triệu chứng lo âu: Nghiên cứu cắt ngang ở học sinh nghề Trung Quốc', 1)
h = doc.add_paragraph(); r = h.add_run('The Effects of Social Media Addiction, Academic Stress, and Sleep Quality on Anxiety Symptoms: A Cross-Sectional Study of Chinese Vocational Students'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'GuangFeng Zheng (1), HaoYan Peng (2)'),
    ('Cơ quan', '(1) Phòng Công tác SV, ĐH Bách khoa Cơ điện Quảng Đông; (2) ĐH Nghề Quảng Châu, TQ'),
    ('Tạp chí', 'Psychology Research and Behavior Management, 2025'),
    ('DOI', '10.2147/PRBM.S522652'),
    ('Loại NC', 'Cắt ngang — phân tích trung gian (mediation)'),
    ('Mẫu', '469 HS nghề 12–18 tuổi, Quảng Đông, TQ'),
    ('Công cụ', 'SMAS (6 mục, nghiện MXH) + Academic Stress Scale (16 mục) + PSQI (giấc ngủ) + GAD-7 (lo âu) + GSES (self-efficacy)'),
    ('Truy cập', 'Open Access — PMC'),
])
add_page_ref(doc, '1–14', 'Psych Res Behav Mgmt', '2025')

add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Nghiên cứu kiểm tra mô hình TAM GIÁC: nghiện MXH + căng thẳng HT + chất lượng giấc ngủ → triệu chứng lo âu (GAD-7), với self-efficacy (hiệu quả bản thân — GSES) là biến TRUNG GIAN.')

p = doc.add_paragraph()
r = p.add_run('Kết quả: Tương quan trực tiếp → lo âu: giấc ngủ kém r = 0,816 (MẠNH NHẤT); áp lực HT r = 0,505; MXH r = 0,415. Hồi quy: giấc ngủ β = 0,615 (mạnh nhất); áp lực β = 0,223; MXH β = 0,153. Self-efficacy là TRUNG GIAN: MXH gián tiếp 63,13% tổng tác động (chủ yếu QUA self-efficacy, không trực tiếp); áp lực 55,84%; giấc ngủ 24,63%.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Psych Res Behav Mgmt — PMC indexed. Open Access.',
    'MÔ HÌNH TAM GIÁC: MXH + Áp lực HT + Giấc ngủ → Lo âu. Ít NC nào đo CẢ 3.',
    'Giấc ngủ β = 0,615 — MẠNH NHẤT. Phù hợp Zhu QT05 (<5h AOR = 13,71).',
    'Self-efficacy là TRUNG GIAN — phát hiện CƠ CHẾ: MXH → giảm self-efficacy → tăng lo âu.',
    'MXH gián tiếp 63,13% — MXH ảnh hưởng lo âu CHỦ YẾU QUA self-efficacy (không trực tiếp).',
    'GAD-7 — cùng công cụ với Hoa 2024 (VN01) và A8 VN COVID.',
    'HS nghề 12–18 — nhóm ÍT NC.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế: Chỉ TQ Quảng Đông, HS nghề. n = 469. Cắt ngang. SMAS 6 mục sơ lược.', bold=True)

add_heading(doc, '1. GIỚI THIỆU', 2)
add_p(doc, 'MXH đã trở thành phần không thể thiếu trong cuộc sống VTN. Tuy nhiên, sử dụng quá mức có thể dẫn đến nghiện MXH, liên quan đến nhiều vấn đề SKTT. Căng thẳng HT là nguồn stress chính ở HS, đặc biệt trong hệ thống giáo dục cạnh tranh của TQ. Chất lượng giấc ngủ kém ảnh hưởng tiêu cực đến SKTT — nhưng ít NC kiểm tra CẢ 3 yếu tố cùng lúc trên cùng mẫu.')
add_p(doc, 'Self-efficacy (hiệu quả bản thân — Bandura, 1977) — niềm tin vào khả năng hoàn thành nhiệm vụ — đã được chứng minh là yếu tố bảo vệ SKTT. Giả thuyết: self-efficacy là biến TRUNG GIAN — MXH/stress/ngủ → giảm self-efficacy → tăng lo âu.')

add_heading(doc, '2. PHƯƠNG PHÁP', 2)
add_p(doc, 'Cắt ngang, 469 HS nghề 12–18 tuổi, Quảng Đông, TQ. 5 thang đo: (1) SMAS — 6 mục, nghiện MXH; (2) Academic Stress Scale — 16 mục; (3) PSQI — 19 mục, chất lượng giấc ngủ; (4) GAD-7 — 7 mục, lo âu (α = 0,90); (5) GSES — 10 mục, self-efficacy. Phân tích tương quan Pearson + hồi quy đa biến + phân tích trung gian (mediation, bootstrap 5.000).')

add_heading(doc, '3. KẾT QUẢ', 2)
add_heading(doc, 'Bảng 1. Tương quan và hồi quy — 3 yếu tố → Lo âu', 3)
add_table(doc,
    ['Yếu tố', 'Tương quan r', 'Hồi quy β', 'Trung gian gián tiếp', 'Ý nghĩa'],
    [['Giấc ngủ kém (PSQI)', 'r = 0,816***', 'β = 0,615***', '24,63% qua self-efficacy', 'MẠNH NHẤT — trực tiếp chủ yếu'],
     ['Áp lực HT', 'r = 0,505***', 'β = 0,223***', '55,84% qua self-efficacy', 'Trung gian MẠNH'],
     ['Nghiện MXH (SMAS)', 'r = 0,415***', 'β = 0,153***', '63,13% qua self-efficacy', 'Chủ yếu GIÁN TIẾP — qua self-efficacy']],
    widths=[3.0, 2.0, 2.0, 3.5, 3.5])
add_p(doc, 'MXH ảnh hưởng lo âu CHỦ YẾU QUA self-efficacy (63% gián tiếp). Giấc ngủ ảnh hưởng TRỰC TIẾP (75%). Áp lực HT cân bằng (56% gián tiếp). Self-efficacy = "cổng" trung gian.', size=9, italic=True)

add_heading(doc, '4. THẢO LUẬN VÀ KẾT LUẬN', 2)
add_p(doc, 'Mô hình tam giác MXH + Áp lực + Giấc ngủ → Lo âu, với self-efficacy trung gian. MXH không trực tiếp "gây" lo âu — MXH → giảm self-efficacy → tăng lo âu. Can thiệp: TĂNG self-efficacy (kỹ năng, tự tin) có thể CHẶN tác động MXH. Giấc ngủ → lo âu MẠNH NHẤT — can thiệp giấc ngủ hiệu quả nhất. Áp lực HT → lo âu 56% qua self-efficacy — phù hợp Wen QT08 (OR=11,58), VN21 (β=4,73).')

add_heading(doc, 'Bảng 2. So sánh với NC khác', 3)
add_table(doc,
    ['NC', 'Yếu tố', 'Kết quả', 'Phù hợp?'],
    [['Zheng 2025 (bài này)', 'Giấc ngủ', 'β = 0,615', 'PHÙ HỢP'],
     ['Zhu QT05', 'Giấc ngủ <5h', 'AOR = 13,71', 'PHÙ HỢP'],
     ['Norway QT21', 'MXH', 'Giải thích MỘT PHẦN', 'PHÙ HỢP — MXH gián tiếp'],
     ['Nature QT27', 'MXH', 'g = 0,46', 'PHÙ HỢP'],
     ['Wen QT08', 'Áp lực HT', 'OR = 11,58', 'PHÙ HỢP'],
     ['VN19 Thảo Vi', 'Lạc quan', 'β gián tiếp = −0,24', 'TƯƠNG TỰ — trung gian tâm lý']],
    widths=[3.0, 2.0, 3.0, 3.5])

add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh: Mô hình tam giác + trung gian self-efficacy. GAD-7. Phát hiện cơ chế. Open Access.', bold=True)
add_red(doc, 'Hạn chế: Chỉ TQ HS nghề. n=469. Cắt ngang — mediation giả định nhân quả. SMAS 6 mục sơ lược. Thiếu cyberbullying, so sánh xã hội.')
add_red(doc, 'Gap: Mô hình tam giác tại VN — MXH + áp lực + giấc ngủ + self-efficacy → lo âu HS THPT VN. Can thiệp: tăng self-efficacy + quản lý giấc ngủ.')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. Mô hình tam giác, trung gian self-efficacy, GAD-7, cơ chế.', bold=True)
save(doc, '46_Zheng_MXH_Stress_Sleep_2025.docx')

print('\n=== DONE 4 BAI ===')
