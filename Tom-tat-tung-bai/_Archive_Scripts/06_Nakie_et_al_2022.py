# -*- coding: utf-8 -*-
"""Bài 6: Nakie et al. 2022 — theo phong cách CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 6 =====
bl('Tóm tắt bài 6', bold=True)

# BƯỚC 1: ĐỊNH DANH
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình nghiên cứu \u00ab Tỷ lệ và các yếu tố liên quan của trầm cảm, lo âu và căng thẳng ở học sinh trung học phổ thông tại Tây Bắc Ethiopia, 2021 \u00bb (Prevalence and Associated Factors of Depression, Anxiety, and Stress among High School Students in Northwest Ethiopia, 2021), do Girum Nakie, Tesfaye Segon, Mamaru Melkam, Getachew Tesfaw Desalegn và Tadele Amare Zeleke (2022), thuộc Khoa Tâm thần, Đại học Gondar và Đại học Metu, Ethiopia, khảo sát trên mẫu 849 học sinh (810 phản hồi, tỷ lệ 96,1%) từ 15 đến 25 tuổi tại 6 trường trung học phổ thông công lập ở Tây Bắc Ethiopia, xuất bản trên BMC Psychiatry, 22:739.')

# BƯỚC 2: TỔNG QUAN PP + KT1
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21 \u2014 Depression Anxiety Stress Scale) kết hợp với Thang đo Ám ảnh Xã hội (SPIN), Thang đo Hỗ trợ Xã hội Oslo-3 và Bài Kiểm tra Sàng lọc ASSIST (Rượu, Hút thuốc và Chất gây nghiện). Nói cách khác, công trình này sử dụng phương pháp sàng lọc cắt ngang với phân tích hồi quy logistic đa biến.', bold=True)

# BƯỚC 3: ĐỊNH NGHĨA (KT5)
bl('DASS-21 là phiên bản rút gọn 21 mục của Thang đo Trầm cảm Lo âu Căng thẳng gốc 42 mục (Lovibond & Lovibond, 1995). Bảng câu hỏi đã được xác thực ở các nước châu Phi với Cronbach alpha: trầm cảm 0,81, lo âu 0,89, căng thẳng 0,78. Pilot test trên 42 học sinh tại Ethiopia cho Cronbach alpha: trầm cảm 0,72, lo âu 0,84, căng thẳng 0,87.')
bl('SPIN (Social Phobia Inventory) là thang đo ám ảnh xã hội với điểm ≥20 gợi ý có ám ảnh xã hội, đã được xác thực ở Nigeria và sử dụng tại Ethiopia.')
bl('ASSIST (Alcohol, Smoking and Substance Involvement Screening Test) là công cụ sàng lọc sử dụng chất gây nghiện do WHO phát triển, phân loại nguy cơ thấp/trung bình/cao.')

# BƯỚC 4: BIỆN MINH (KT4)
bl('Tổng quan tài liệu của chúng tôi cho thấy, mặc dù trầm cảm, lo âu và căng thẳng ở học sinh trung học phổ thông đã được nghiên cứu ở các nước phát triển, không có nghiên cứu nào được công bố ở châu Phi về thanh thiếu niên trung học phổ thông (trước nghiên cứu này). Đây là nghiên cứu tiên phong tại lục địa này.', bold=True)

# BƯỚC 5: LIỆT KÊ DỮ LIỆU
bl('Dữ liệu nhân khẩu học-xã hội bao gồm tuổi, giới tính, khối lớp, nơi xuất thân (nông thôn/thành thị), sắp xếp sống (cả hai cha mẹ/đơn thân/một mình/người thân), trình độ học vấn của cha và mẹ, tình trạng hôn nhân, kết quả học tập. Dữ liệu lâm sàng: ám ảnh xã hội, bệnh mãn tính, tiền sử gia đình bệnh tâm thần, ý tưởng và nỗ lực tự sát, hỗ trợ xã hội. Dữ liệu sử dụng chất: rượu, khat, thuốc lá.')

# BƯỚC 6: MÔ TẢ QUY TRÌNH
bl('Nghiên cứu cắt ngang được thực hiện vào tháng 4/2021 tại 6 trường trung học phổ thông công lập được chọn ngẫu nhiên ở vùng Amhara, Tây Bắc Ethiopia. Lấy mẫu phân tầng đa giai đoạn: chọn ngẫu nhiên trường → phân bổ tỷ lệ theo khối lớp → chọn ngẫu nhiên đơn giản từng học sinh. DASS-21 được dịch sang tiếng Amharic và dịch ngược bởi chuyên gia ngôn ngữ. Pilot test trên 42 học sinh. Phân tích bằng SPSS 25.0: hồi quy logistic nhị phân, P < 0,05 trong phân tích đa biến. Kiểm định Hosmer-Lemeshow để đánh giá phù hợp mô hình.')

# BƯỚC 9: ĐỊNH LƯỢNG — BẢNG
rb('Kết quả nghiên cứu định lượng')

bl('Tỷ lệ trầm cảm, lo âu và căng thẳng theo DASS-21 (N = 810):', bold=True)
tbl(
    ['Tình trạng', 'Tỷ lệ', 'KTC 95%'],
    [
        ['Trầm cảm', '41,4%', '38,0\u201344,8%'],
        ['Lo âu', '66,7%', '66,4\u201366,9%'],
        ['Căng thẳng', '52,2%', '49,1\u201356,0%'],
    ],
    widths=[4.0, 3.0, 4.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan đến trầm cảm (phân tích đa biến):', bold=True)
tbl(
    ['Yếu tố', 'AOR', 'KTC 95%', 'P'],
    [
        ['Nhai khat nguy cơ cao', '5,595', '2,357\u201311,132', '<0,001'],
        ['Ám ảnh xã hội', '1,416', '1,045\u20131,919', '0,025'],
        ['Giới tính nữ', '1,304', '1,006\u20131,849', '0,046'],
    ],
    widths=[4.5, 2.0, 3.0, 2.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan đến lo âu (phân tích đa biến):', bold=True)
tbl(
    ['Yếu tố', 'AOR', 'KTC 95%', 'P'],
    [
        ['Hút thuốc lá nguy cơ cao', '4,777', '1,407\u20137.304', '0,012'],
        ['Bệnh mãn tính đồng mắc', '2,099', '1,045\u20134,218', '0,037'],
        ['Tiền sử bệnh TT gia đình', '1,777', '1,028\u20133,073', '0,040'],
    ],
    widths=[4.5, 2.0, 3.0, 2.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan đến căng thẳng (phân tích đa biến):', bold=True)
tbl(
    ['Yếu tố', 'AOR', 'KTC 95%', 'P'],
    [
        ['Uống rượu nguy cơ cao', '1,828', '1,012\u20133,303', '<0,05'],
        ['Hỗ trợ xã hội kém', '1,739', '1,203\u20132,515', '<0,05'],
        ['Cư trú nông thôn', '1,395', '1,010\u20131,925', '<0,05'],
    ],
    widths=[4.5, 2.0, 3.0, 2.0]
)
doc.add_paragraph()

# BƯỚC 10: ĐỐI CHIẾU (KT3)
bh3('Khác biệt giới tính : Nữ giới luôn được báo cáo có tỷ lệ trầm cảm cao hơn nam giới trong y văn quốc tế (Salk và cộng sự, 2017). Thực tế, nghiên cứu tại các quốc gia đang phát triển cũng xác nhận xu hướng này. Chúng tôi xác nhận trong mẫu học sinh Ethiopia rằng nữ có nguy cơ trầm cảm cao gấp 1,3 lần so với nam (AOR = 1,304, P = 0,046) \u2014 phù hợp với xu hướng quốc tế nhưng mức chênh lệch khiêm tốn hơn so với các nghiên cứu phương Tây.')

# BƯỚC 11+12: NHẬN XÉT
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về vai trò của chất gây nghiện và yếu tố xã hội trong bối cảnh châu Phi')
bl('*Khat và thuốc lá.* Nổi bật nhất là mối liên hệ cực mạnh giữa nhai khat nguy cơ cao và trầm cảm (AOR = 5,595) cũng như hút thuốc lá nguy cơ cao và lo âu (AOR = 4,777). Đây là hai yếu tố nguy cơ đặc thù văn hóa Đông Phi mà các nghiên cứu tại châu Á và phương Tây không đề cập. Khat (Catha edulis) là chất kích thích phổ biến ở vùng Sừng châu Phi, gây tác động trên hệ thần kinh trung ương tương tự amphetamine.')
bl('*Hỗ trợ xã hội.* Học sinh có hỗ trợ xã hội kém (Oslo-3 điểm 3\u20138) có nguy cơ căng thẳng cao gấp 1,7 lần (AOR = 1,739). Phát hiện này tương đồng với kết luận của Pham và cộng sự (2024) tại Việt Nam về vai trò bảo vệ của chăm sóc cảm xúc (beta = \u20130,40) và của NSCH (2020) về thanh thiếu niên có SKTT gặp khó khăn kết bạn gấp 10 lần.')
bl('*Cư trú nông thôn.* Học sinh từ nông thôn (70,5% mẫu) có nguy cơ căng thẳng cao gấp 1,4 lần so với thành thị (AOR = 1,395). Điều này tương tự với Xu và cộng sự (2021) tại Trung Quốc, nơi nông thôn có tỷ lệ lo âu cao hơn (12,80%) so với thành thị.')

# BƯỚC 13: KẾT LUẬN (Kiến trúc B)
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy tỷ lệ rất cao của trầm cảm (41,4%), lo âu (66,7%) và căng thẳng (52,2%) ở học sinh trung học phổ thông Tây Bắc Ethiopia \u2014 nghiên cứu tiên phong tại châu Phi, gợi ý rằng sức khỏe tâm thần thanh thiếu niên tại các nước đang phát triển cần được chú ý khẩn cấp. Các yếu tố nguy cơ đặc thù văn hóa \u2014 nhai khat (AOR = 5,6), hút thuốc lá (AOR = 4,8) \u2014 và yếu tố xã hội \u2014 hỗ trợ xã hội kém, cư trú nông thôn \u2014 nhấn mạnh nhu cầu can thiệp nhạy cảm bối cảnh. Trong bối cảnh vùng nông thôn Đông Phi như nghiên cứu hiện tại, nhu cầu mở rộng dịch vụ sức khỏe tâm thần và tăng cường tư vấn tại trường học là thiết yếu.', bold=True)

# PHẢN BIỆN
rh2('Phản biện')
bl('Điểm mạnh nổi bật: xuất bản trên BMC Psychiatry (tạp chí Q1, PubMed-indexed), sử dụng phân tích hồi quy đa biến (AOR) \u2014 chuẩn phương pháp luận vượt trội so với hầu hết nghiên cứu trong tập hợp 11 bài. Lấy mẫu phân tầng đa giai đoạn từ 6 trường, pilot test trên 42 học sinh, Cronbach alpha được báo cáo đầy đủ. Tuy nhiên, tỷ lệ lo âu 66,7% là rất cao \u2014 DASS-21 là thang sàng lọc với ngưỡng cắt thấp (≥8 cho lo âu), gây câu hỏi về tỷ lệ dương tính giả tương tự như Alharbi và cộng sự (2019). Thiết kế cắt ngang không cho phép suy luận nhân quả \u2014 mối quan hệ giữa khat/thuốc lá và DAS có thể là hai chiều. Mẫu bao gồm học sinh từ 15 đến 25 tuổi (33,7% ≥20 tuổi) \u2014 rộng hơn phạm vi thanh thiếu niên thông thường.')

# HƯỚNG NC
rh2('Hướng nghiên cứu tiếp theo')
bl('Nghiên cứu dọc để xác định chiều hướng nhân quả giữa sử dụng chất gây nghiện (khat, thuốc lá, rượu) và trầm cảm/lo âu/căng thẳng. Nghiên cứu can thiệp giảm sử dụng khat ở thanh thiếu niên và đánh giá tác động lên sức khỏe tâm thần. So sánh DASS-21 với công cụ chẩn đoán (DISC-5, MINI) trên cùng mẫu để ước tính khoảng cách sàng lọc\u2013chẩn đoán tại Ethiopia. Mở rộng ra các vùng khác của Ethiopia và các nước Đông Phi để đánh giá tính khái quát.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50 Cao. BMC Psychiatry Q1, phân tích AOR đa biến, lấy mẫu phân tầng, Cronbach alpha đầy đủ, nghiên cứu tiên phong tại châu Phi.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('06_Nakie_et_al_2022.docx')
print('06_Nakie_et_al_2022.docx OK')
