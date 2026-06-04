# -*- coding: utf-8 -*-
"""Dịch batch 6 bài mới: B5, B3, B6, B11, A7, A12"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                t += ' ' + cell.text
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

# ========== B5 — UK School Interventions 2025 ==========
print('B5 UK school...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1080/09638237.2025.2512332', size=10)
add_heading(doc, 'Can thi\u1ec7p t\u1ea1i tr\u01b0\u1eddng cho tr\u1ea7m c\u1ea3m v\u00e0 lo \u00e2u \u1edf V\u01b0\u01a1ng qu\u1ed1c Anh', 1)
h = doc.add_paragraph(); r = h.add_run('School based interventions for depression and anxiety in UK'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('T\u1ea1p ch\u00ed', 'Journal of Mental Health (Q1, IF \u2248 3,5) \u2014 Taylor & Francis'),
    ('Xu\u1ea5t b\u1ea3n', '2025, 6 trang'),
    ('DOI', '10.1080/09638237.2025.2512332'),
    ('Lo\u1ea1i NC', 'T\u1ed5ng quan h\u1ec7 th\u1ed1ng (SR) \u2014 can thi\u1ec7p tr\u01b0\u1eddng h\u1ecdc cho tr\u1ea7m c\u1ea3m/lo \u00e2u t\u1ea1i UK'),
    ('Ph\u1ea1m vi', 'UK \u2014 h\u1ec7 th\u1ed1ng gi\u00e1o d\u1ee5c Anh'),
])
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in ['J Mental Health Q1. Can thi\u1ec7p tr\u01b0\u1eddng UK \u2014 b\u1ed1i c\u1ea3nh NHS.', 'T\u1ed5ng quan c\u00e1c ch\u01b0\u01a1ng tr\u00ecnh \u0111\u00e3 tri\u1ec3n khai t\u1ea1i UK.', 'Ph\u00f9 h\u1ee3p \u0111\u1ec1 t\u00e0i: m\u00f4 h\u00ecnh can thi\u1ec7p tr\u01b0\u1eddng t\u1eeb n\u01b0\u1edbc ph\u00e1t tri\u1ec3n \u2014 tham kh\u1ea3o cho VN.']:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf: Ch\u1ec9 UK. H\u1ec7 th\u1ed1ng gi\u00e1o d\u1ee5c kh\u00e1c VN.', bold=True)
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N')
add_red(doc, 'J Mental Health Q1. UK school \u2014 NHS context. C\u1ea7n th\u00edch \u1ee9ng cho VN (B7 CA-CBT \u0110NA). Ph\u00f9 h\u1ee3p B8 Sri Lanka RCT (GV cung c\u1ea5p). Gap: VN ch\u01b0a c\u00f3 ch\u01b0\u01a1ng tr\u00ecnh t\u01b0\u01a1ng t\u1ef1.')
add_p(doc, '\u0110\u00e1nh gi\u00e1: \u2b50\u2b50\u2b50 Trung b\u00ecnh. SR UK, ng\u1eafn 6p, ch\u1ec9 UK.', bold=True)
save(doc, '48_UK_School_Interventions_2025.docx')

# ========== B3 — JAMA Mobile App CBT Anxiety RCT 2024 ==========
print('B3 JAMA App CBT...')
doc = create_doc()
add_p(doc, 'Link: https://jamanetwork.com/journals/jamanetworkopen/fullarticle/2822451', size=10)
add_heading(doc, 'Hi\u1ec7u qu\u1ea3 can thi\u1ec7p d\u1ef1a tr\u00ean \u1ee9ng d\u1ee5ng di \u0111\u1ed9ng cho thanh ni\u00ean c\u00f3 r\u1ed1i lo\u1ea1n lo \u00e2u: Th\u1eed nghi\u1ec7m l\u00e2m s\u00e0ng ng\u1eabu nhi\u00ean', 1)
h = doc.add_paragraph(); r = h.add_run('Efficacy of a Mobile App-Based Intervention for Young Adults With Anxiety Disorders: A Randomized Clinical Trial'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('T\u00e1c gi\u1ea3', 'Jennifer N. Bress et al.'),
    ('T\u1ea1p ch\u00ed', 'JAMA Network Open (Q1, IF \u2248 13,8)'),
    ('Xu\u1ea5t b\u1ea3n', '2024, 15 trang'),
    ('DOI', '10.1001/jamanetworkopen.2024.28757'),
    ('Lo\u1ea1i NC', 'RCT \u2014 app CBT t\u1ef1 h\u01b0\u1edbng d\u1eabn cho thanh ni\u00ean lo \u00e2u'),
    ('M\u1eabu', 'Thanh ni\u00ean 18\u201325 tu\u1ed5i v\u1edbi r\u1ed1i lo\u1ea1n lo \u00e2u'),
    ('C\u00f4ng c\u1ee5', 'GAD-7, OASIS'),
])
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'JAMA NetOpen Q1 IF \u2248 13,8 \u2014 t\u1ea1p ch\u00ed UY T\u00cdN C\u1ef0C CAO.',
    'RCT app CBT t\u1ef1 h\u01b0\u1edbng d\u1eabn \u2014 b\u1eb1ng ch\u1ee9ng nh\u00e2n qu\u1ea3.',
    'App di \u0111\u1ed9ng \u2014 ti\u1ebfp c\u1eadn d\u1ec5, chi ph\u00ed th\u1ea5p, kh\u00f4ng c\u1ea7n chuy\u00ean gia.',
    'Ph\u00f9 h\u1ee3p B2 JMIR (DMHI SAD g=0,508) v\u00e0 B9 NMA (iCBT h\u1ea1ng 1 SUCRA 71,2%).',
    'Thanh ni\u00ean 18\u201325 \u2014 r\u1ed9ng h\u01a1n VTN nh\u01b0ng \u00e1p d\u1ee5ng \u0111\u01b0\u1ee3c.',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf: Ch\u1ec9 thanh ni\u00ean (18\u201325), kh\u00f4ng VTN. Ch\u1ec9 M\u1ef9.', bold=True)
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N')
add_red(doc, 'JAMA Q1 IF=13,8. RCT app CBT. Ph\u00f9 h\u1ee3p B2 JMIR + B9 NMA: iCBT hi\u1ec7u qu\u1ea3 cho lo \u00e2u. H\u1ea1n ch\u1ebf: ch\u1ec9 M\u1ef9, 18\u201325 tu\u1ed5i. Gap: app CBT ti\u1ebfng Vi\u1ec7t cho VTN VN.')
add_p(doc, '\u0110\u00e1nh gi\u00e1: \u2b50\u2b50\u2b50\u2b50 Cao. JAMA Q1, RCT, app CBT.', bold=True)
save(doc, '49_JAMA_App_CBT_2024.docx')

# ========== B6 — School Resilience SR+MA 2025 ==========
print('B6 Resilience...')
doc = create_doc()
add_p(doc, 'Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC12127306/', size=10)
add_heading(doc, 'Can thi\u1ec7p t\u1ea1i tr\u01b0\u1eddng cho kh\u1ea3 n\u0103ng ph\u1ee5c h\u1ed3i \u1edf tr\u1ebb em v\u00e0 thanh thi\u1ebfu ni\u00ean: T\u1ed5ng quan h\u1ec7 th\u1ed1ng v\u00e0 ph\u00e2n t\u00edch t\u1ed5ng h\u1ee3p c\u00e1c th\u1eed nghi\u1ec7m ng\u1eabu nhi\u00ean c\u00f3 \u0111\u1ed1i ch\u1ee9ng', 1)
h = doc.add_paragraph(); r = h.add_run('School-based interventions for resilience in children and adolescents: a systematic review and meta-analysis of randomized controlled trials'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('T\u00e1c gi\u1ea3', 'Chenyi Cao et al.'),
    ('T\u1ea1p ch\u00ed', 'Frontiers in Psychiatry (Q1, IF \u2248 4,7)'),
    ('Xu\u1ea5t b\u1ea3n', '2025, 15 trang'),
    ('DOI', '10.3389/fpsyt.2025.1594658'),
    ('Lo\u1ea1i NC', 'T\u1ed5ng quan h\u1ec7 th\u1ed1ng + Ph\u00e2n t\u00edch t\u1ed5ng h\u1ee3p (SR + MA) c\u00e1c RCT'),
    ('M\u1eabu', 'Nhi\u1ec1u RCT v\u1ec1 can thi\u1ec7p resilience t\u1ea1i tr\u01b0\u1eddng'),
    ('Truy c\u1eadp', 'Open Access \u2014 PMC/Frontiers'),
])
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'Frontiers in Psychiatry Q1 IF \u2248 4,7. Open Access.',
    'SR + MA c\u00e1c RCT \u2014 b\u1eb1ng ch\u1ee9ng m\u1ea1nh v\u1ec1 can thi\u1ec7p RESILIENCE t\u1ea1i tr\u01b0\u1eddng.',
    'Resilience = kh\u1ea3 n\u0103ng ph\u1ee5c h\u1ed3i \u2014 y\u1ebfu t\u1ed1 B\u1ea2O V\u1ec6 quan tr\u1ecdng.',
    'Ireland QT32: resilience quan tr\u1ecdng h\u01a1n theo th\u1eddi gian. \u00dac QT25: flourishing gi\u1ea3m.',
    'Ph\u00f9 h\u1ee3p \u0111\u1ec1 c\u01b0\u01a1ng giai \u0111o\u1ea1n 2: CBT + PE + RESILIENCE t\u1ea1i tr\u01b0\u1eddng.',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf: \u0110a s\u1ed1 RCT ph\u01b0\u01a1ng T\u00e2y. Ch\u01b0a r\u00f5 hi\u1ec7u qu\u1ea3 \u1edf ch\u00e2u \u00c1/LMIC.', bold=True)
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N')
add_red(doc, 'Frontiers Q1. SR+MA RCT resilience. Ph\u00f9 h\u1ee3p Ireland QT32 (resilience + t\u1ef1 tr\u1ecdng). B8 Sri Lanka RCT c\u0169ng t\u0103ng t\u1ef1 tr\u1ecdng (\u03b2=0,811). Gap: RCT resilience t\u1ea1i VN.')
add_p(doc, '\u0110\u00e1nh gi\u00e1: \u2b50\u2b50\u2b50\u2b50 Cao. SR+MA RCTs, resilience, Frontiers Q1.', bold=True)
save(doc, '50_Resilience_School_MA_2025.docx')

# ========== B11 — Japan iCBT SAD 2024 ==========
print('B11 Japan iCBT...')
doc = create_doc()
add_p(doc, 'Link: https://pediatrics.jmir.org/2024/1/e55786', size=10)
add_heading(doc, 'Hi\u1ec7u qu\u1ea3 li\u1ec7u ph\u00e1p nh\u1eadn th\u1ee9c\u2013h\u00e0nh vi qua internet kh\u00f4ng h\u01b0\u1edbng d\u1eabn cho r\u1ed1i lo\u1ea1n lo \u00e2u x\u00e3 h\u1ed9i d\u01b0\u1edbi ng\u01b0\u1ee1ng \u1edf thanh thi\u1ebfu ni\u00ean v\u00e0 thanh ni\u00ean: Th\u1eed nghi\u1ec7m \u0111a trung t\u00e2m ng\u1eabu nhi\u00ean c\u00f3 \u0111\u1ed1i ch\u1ee9ng', 1)
h = doc.add_paragraph(); r = h.add_run('Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: Multicenter Randomized Controlled Trial'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('T\u1ea1p ch\u00ed', 'JMIR Pediatrics and Parenting (Q2, IF \u2248 2,5)'),
    ('Xu\u1ea5t b\u1ea3n', '2024, 13 trang'),
    ('DOI', '10.2196/55786'),
    ('Lo\u1ea1i NC', 'RCT \u0111a trung t\u00e2m \u2014 iCBT KH\u00d4NG h\u01b0\u1edbng d\u1eabn cho SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng'),
    ('M\u1eabu', 'VTN v\u00e0 thanh ni\u00ean Nh\u1eadt B\u1ea3n v\u1edbi SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng'),
    ('C\u00f4ng c\u1ee5', 'LSAS (Liebowitz Social Anxiety Scale)'),
])
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'JMIR Pediatrics Q2. RCT \u0111a trung t\u00e2m Nh\u1eadt B\u1ea3n.',
    'iCBT KH\u00d4NG h\u01b0\u1edbng d\u1eabn \u2014 t\u1ef1 h\u1ecdc ho\u00e0n to\u00e0n. Chi ph\u00ed th\u1ea5p nh\u1ea5t.',
    'SAD D\u01af\u1edaI NG\u01af\u1ee0NG \u2014 can thi\u1ec7p S\u1edaM tr\u01b0\u1edbc khi th\u00e0nh r\u1ed1i lo\u1ea1n \u0111\u1ea7y \u0111\u1ee7.',
    'Nh\u1eadt B\u1ea3n \u2014 \u0110\u00f4ng B\u1eafc \u00c1, v\u0103n h\u00f3a g\u1ea7n VN h\u01a1n ph\u01b0\u01a1ng T\u00e2y.',
    'Ph\u00f9 h\u1ee3p B9 NMA (iCBT SUCRA 71,2% h\u1ea1ng 1) v\u00e0 B2 JMIR (DMHI g=0,508).',
    'Nh\u01b0ng B2 cho th\u1ea5y c\u00f3 h\u01b0\u1edbng d\u1eabn (g=0,825) > kh\u00f4ng h\u01b0\u1edbng d\u1eabn (~0,3\u20130,4).',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf: Ch\u1ec9 Nh\u1eadt. SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng (kh\u00f4ng ph\u1ea3i \u0111\u1ea7y \u0111\u1ee7). Kh\u00f4ng h\u01b0\u1edbng d\u1eabn \u2014 hi\u1ec7u qu\u1ea3 c\u00f3 th\u1ec3 h\u1ea1n ch\u1ebf.', bold=True)
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N')
add_red(doc, 'JMIR Pediatrics. RCT Nh\u1eadt B\u1ea3n \u2014 \u0110BA g\u1ea7n VN. iCBT t\u1ef1 h\u1ecdc \u2014 chi ph\u00ed th\u1ea5p nh\u1ea5t. SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng \u2014 can thi\u1ec7p s\u1edbm. Nh\u01b0ng: B2 JMIR cho th\u1ea5y c\u00f3 h\u01b0\u1edbng d\u1eabn (g=0,825) v\u01b0\u1ee3t tr\u1ed9i kh\u00f4ng h\u01b0\u1edbng d\u1eabn. Gap: iCBT VN c\u1ea7n c\u00f3 GV/mentor h\u01b0\u1edbng d\u1eabn.')
add_p(doc, '\u0110\u00e1nh gi\u00e1: \u2b50\u2b50\u2b50\u2b50 Cao. RCT iCBT Nh\u1eadt, SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng, \u0110BA.', bold=True)
save(doc, '51_Japan_iCBT_SAD_2024.docx')

# ========== A7 — Academic Stress Interventions SR 2024 ==========
print('A7 Academic Stress Interventions...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1007/s10578-024-01667-5', size=10)
add_heading(doc, 'Can thi\u1ec7p c\u0103ng th\u1eb3ng h\u1ecdc t\u1eadp \u1edf tr\u01b0\u1eddng trung h\u1ecdc: T\u1ed5ng quan t\u00e0i li\u1ec7u h\u1ec7 th\u1ed1ng', 1)
h = doc.add_paragraph(); r = h.add_run('Academic Stress Interventions in High Schools: A Systematic Literature Review'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('T\u1ea1p ch\u00ed', 'Child Psychiatry & Human Development (Q1, IF \u2248 3,3) \u2014 Springer'),
    ('Xu\u1ea5t b\u1ea3n', '2025, Vol. 56, pp. 1836\u20131869, 34 trang'),
    ('DOI', '10.1007/s10578-024-01667-5'),
    ('Lo\u1ea1i NC', 'T\u1ed5ng quan t\u00e0i li\u1ec7u h\u1ec7 th\u1ed1ng (SLR) \u2014 can thi\u1ec7p stress HT'),
    ('Ph\u1ea1m vi', 'HS trung h\u1ecdc (secondary school) to\u00e0n c\u1ea7u'),
])
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'Child Psychiatry Q1 IF \u2248 3,3. Springer. 34 trang \u2014 to\u00e0n di\u1ec7n.',
    'T\u1ed4NG QUAN DUY NH\u1ea4T v\u1ec1 can thi\u1ec7p STRESS H\u1eccC T\u1eacP ri\u00eang \u1edf HS trung h\u1ecdc.',
    'Stress HT = y\u1ebfu t\u1ed1 M\u1ea0NH NH\u1ea4T t\u1ea1i VN: VN21 \u03b2=4,73; Wen QT08 OR=11,58; Norway QT21 gi\u1ea3i th\u00edch ch\u00ednh.',
    'R\u1ea4T PH\u00d9 H\u1ee2P \u0111\u1ec1 c\u01b0\u01a1ng giai \u0111o\u1ea1n 2 \u2014 can thi\u1ec7p gi\u1ea3m stress HT t\u1ea1i tr\u01b0\u1eddng VN.',
    'UNICEF VN22: 47% HS h\u1ecdc th\u00eam >3h/tu\u1ea7n, 15% >9h/tu\u1ea7n \u2014 c\u1ea7n can thi\u1ec7p.',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf: \u0110a s\u1ed1 NC ph\u01b0\u01a1ng T\u00e2y. 34 trang d\u00e0i \u2014 c\u1ea7n t\u1eadp trung k\u1ebft qu\u1ea3.', bold=True)
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N')
add_red(doc, 'Child Psychiatry Q1. SLR can thi\u1ec7p stress HT \u2014 DUY NH\u1ea4T. Stress HT = y\u1ebfu t\u1ed1 m\u1ea1nh nh\u1ea5t VN (VN21, QT08, QT21, UNICEF VN22). Ph\u00f9 h\u1ee3p \u0111\u1ec1 c\u01b0\u01a1ng GD 2. H\u1ea1n ch\u1ebf: \u0111a s\u1ed1 ph\u01b0\u01a1ng T\u00e2y, thi\u1ebfu LMIC/ch\u00e2u \u00c1. Gap: can thi\u1ec7p stress HT ri\u00eang t\u1ea1i VN (kh\u00e1c CBT chung).')
add_p(doc, '\u0110\u00e1nh gi\u00e1: \u2b50\u2b50\u2b50\u2b50 Cao. SLR duy nh\u1ea5t, Springer Q1, r\u1ea5t ph\u00f9 h\u1ee3p \u0111\u1ec1 t\u00e0i.', bold=True)
save(doc, '52_AcademicStress_Interventions_SR_2025.docx')

# ========== A12 — Dinh 2021 School Factors VN ==========
print('A12 Dinh School Factors VN...')
doc = create_doc()
add_p(doc, 'Link: https://www.researchgate.net/publication/349382974', size=10)
add_heading(doc, 'Y\u1ebfu t\u1ed1 tr\u01b0\u1eddng h\u1ecdc g\u00e2y lo \u00e2u \u1edf thanh thi\u1ebfu ni\u00ean Vi\u1ec7t Nam t\u1ea1i c\u00e1c tr\u01b0\u1eddng trung h\u1ecdc c\u01a1 s\u1edf', 1)
h = doc.add_paragraph(); r = h.add_run('School Factors Causing Vietnamese Adolescents\' Anxiety in Secondary Schools'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('Ti\u00eau \u0111\u1ec1 g\u1ed1c', 'School Factors Causing Vietnamese Adolescents\' Anxiety in Secondary Schools'),
    ('T\u00e1c gi\u1ea3', 'Dinh HVT, Do LHT, Phan MHT'),
    ('T\u1ea1p ch\u00ed', 'Psychology and Education, 2021, 58(1): 883\u2013894'),
    ('Lo\u1ea1i NC', 'NC g\u1ed1c \u2014 y\u1ebfu t\u1ed1 tr\u01b0\u1eddng h\u1ecdc \u2192 lo \u00e2u \u1edf HS THCS VN'),
    ('M\u1eabu', 'HS THCS Vi\u1ec7t Nam'),
    ('L\u01b0u \u00fd', 'B\u00e0i n\u00e0y \u0111\u01b0\u1ee3c tr\u00edch d\u1eabn trong VN21 (Tr\u1ea7n Th\u1ea3o Vi 2024) \u2014 ref #33'),
])
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'NC G\u1ed0C t\u1ea1i VN v\u1ec1 y\u1ebfu t\u1ed1 TR\u01af\u1edcNG H\u1eccC g\u00e2y lo \u00e2u \u1edf HS THCS.',
    '\u0110I\u1ec2M S\u1ed0 + THI C\u1eec = ngu\u1ed3n lo \u00e2u CH\u00cdNH. N\u1eef > nam.',
    'L\u1edbp 6\u20138 lo \u00e2u CAO H\u01a0N l\u1edbp 9 v\u1ec1 y\u1ebfu t\u1ed1 b\u1ea1n b\u00e8/GV/tr\u01b0\u1eddng.',
    '\u0110\u01b0\u1ee3c tr\u00edch d\u1eabn b\u1edfi VN21 (Tr\u1ea7n Th\u1ea3o Vi) \u2014 \u0111\u00e3 x\u00e1c nh\u1eadn.',
    'Ph\u00f9 h\u1ee3p UNICEF VN22 (HS kh\u00f4ng tho\u1ea3i m\u00e1i t\u00ecm GV), Norway QT21 (b\u1ea5t m\u00e3n tr\u01b0\u1eddng gi\u1ea3i th\u00edch ch\u00ednh).',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf: TC Psychology & Education \u2014 kh\u00f4ng PubMed/Scopus. N\u0103m 2021.', bold=True)

add_heading(doc, 'K\u1ebeT QU\u1ea2 CH\u00cdNH', 2)
add_heading(doc, 'B\u1ea3ng 1. Y\u1ebfu t\u1ed1 tr\u01b0\u1eddng h\u1ecdc g\u00e2y lo \u00e2u HS THCS VN', 3)
add_table(doc,
    ['Y\u1ebfu t\u1ed1', 'M\u1ee9c \u0111\u1ed9 lo \u00e2u', 'Gi\u1edbi t\u00ednh', 'Kh\u1ed1i l\u1edbp'],
    [['\u0110i\u1ec3m s\u1ed1 + thi c\u1eed', 'CAO NH\u1ea4T', 'N\u1eef > Nam', ''],
     ['Giao ti\u1ebfp v\u1edbi b\u1ea1n b\u00e8', 'Trung b\u00ecnh', '', 'L\u1edbp 6\u20138 > L\u1edbp 9'],
     ['Giao ti\u1ebfp v\u1edbi GV', 'Trung b\u00ecnh', '', 'L\u1edbp 6\u20138 > L\u1edbp 9'],
     ['C\u01a1 s\u1edf v\u1eadt ch\u1ea5t tr\u01b0\u1eddng', 'Th\u1ea5p', '', 'L\u1edbp 6\u20138 > L\u1edbp 9'],
     ['An to\u00e0n tr\u01b0\u1eddng', 'Th\u1ea5p', '', '']],
    widths=[4.0, 2.5, 2.5, 3.0])
add_p(doc, '\u0110i\u1ec3m s\u1ed1 + thi c\u1eed g\u00e2y lo \u00e2u NHI\u1ec0U NH\u1ea4T. N\u1eef > nam \u2014 ph\u00f9 h\u1ee3p to\u00e0n c\u1ea7u (Islam QT31: AOR=1,51). L\u1edbp 6\u20138 > l\u1edbp 9 v\u1ec1 y\u1ebfu t\u1ed1 XH (b\u1ea1n b\u00e8, GV) \u2014 c\u00f3 th\u1ec3 do l\u1edbp 9 \u0111\u00e3 th\u00edch nghi.', size=9, italic=True)

add_heading(doc, 'B\u1ea3ng 2. So s\u00e1nh v\u1edbi NC kh\u00e1c', 3)
add_table(doc,
    ['NC', 'Y\u1ebfu t\u1ed1', 'Ph\u00f9 h\u1ee3p?'],
    [['Dinh 2021 (b\u00e0i n\u00e0y)', '\u0110i\u1ec3m s\u1ed1 + thi c\u1eed = ch\u00ednh', '\u2014'],
     ['VN21 Tr\u1ea7n Th\u1ea3o Vi', 'H\u1ecdc th\u00eam \u03b2=4,73 (ESSA)', 'PH\u00d9 H\u1ee2P \u2014 \u00e1p l\u1ef1c HT'],
     ['Wen QT08', '\u00c1p l\u1ef1c OR=11,58', 'PH\u00d9 H\u1ee2P'],
     ['Norway QT21', 'B\u1ea5t m\u00e3n tr\u01b0\u1eddng gi\u1ea3i th\u00edch ch\u00ednh', 'PH\u00d9 H\u1ee2P'],
     ['UNICEF VN22', 'HS kh\u00f4ng tho\u1ea3i m\u00e1i t\u00ecm GV', 'PH\u00d9 H\u1ee2P']],
    widths=[3.5, 4.0, 3.0])

add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N')
add_red(doc, 'NC g\u1ed1c VN v\u1ec1 y\u1ebfu t\u1ed1 tr\u01b0\u1eddng h\u1ecdc \u2192 lo \u00e2u HS THCS. \u0110i\u1ec3m s\u1ed1 + thi c\u1eed = ch\u00ednh. N\u1eef > nam. L\u1edbp 6\u20138 > l\u1edbp 9 v\u1ec1 XH. \u0110\u01b0\u1ee3c tr\u00edch d\u1eabn b\u1edfi VN21. H\u1ea1n ch\u1ebf: TC kh\u00f4ng PubMed, n\u0103m 2021. Gap: NC c\u1eadp nh\u1eadt 2024\u20132025 v\u1edbi ESSA + GAD-7 t\u1ea1i VN.')
add_p(doc, '\u0110\u00e1nh gi\u00e1: \u2b50\u2b50\u2b50 Trung b\u00ecnh. NC g\u1ed1c VN, \u0111\u01b0\u1ee3c tr\u00edch d\u1eabn, nh\u01b0ng TC kh\u00f4ng PubMed.', bold=True)
save(doc, '53_Dinh_SchoolFactors_VN_2021.docx')

print('\n=== DONE 6 BAI ===')
