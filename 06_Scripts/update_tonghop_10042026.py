# -*- coding: utf-8 -*-
"""
Update Tổng hợp liên bài (09042026 → 10042026)
- Mở doc cũ (đã có 11 sections, 12 tables, 35 bài + 9 bài can thiệp/VN)
- Thêm sections 12-15: bài 45-58 (14 bài mới)
"""
import sys, os, shutil, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Tổng hợp liên bài báo - Lo âu HS - 09042026.docx')
DST = os.path.join(ROOT, 'Tổng hợp liên bài báo - Lo âu HS - 10042026.docx')

shutil.copy(SRC, DST)
doc = Document(DST)

# helpers
def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    tcW.append(w)

def H(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(text, bold=False, italic=False, color=None, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color
    return p

def table(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# Add a page break before the new section
doc.add_page_break()

# ============================================================
# SECTION 12 — UPDATE 10/04/2026
# ============================================================
H('12. CẬP NHẬT 10/04/2026 — 14 bài bổ sung (45-58)', level=2)
P('Sau phiên 10/04/2026, dự án đã bổ sung 14 bài mới: 10 bài có PDF đầy đủ + 4 bài abstract-only (paywall). Tổng số bài hiện nay: 58 (gồm 35 bài cũ + 9 bài VN/can thiệp đã thêm trong phiên 09/04 + 14 bài mới phiên này). Phần này tổng hợp các phát hiện chính + đối chiếu liên bài cho 14 bài mới.', italic=True)

H('Bảng 13. Tổng hợp 14 bài mới phiên 10/04/2026', level=3)
table(
    ['#', 'Bài', 'Loại NC', 'n / Mẫu', 'Phát hiện chính'],
    [
        ['45', 'Hải Phòng 2024 (A10)', 'Cắt ngang VN', '420 HS THPT', 'Lo âu 39,3% (DASS-21); nữ > nam'],
        ['46', 'Zheng 2025 TQ MXH (A6)', 'Cắt ngang + mediation', '469 HS nghề', 'Giấc ngủ β=0,615 (mạnh nhất); MXH gián tiếp 63% qua self-efficacy'],
        ['47', 'Long An 2025 (A11)', 'Cắt ngang VN', 'HS THPT', 'Lo âu 57,2% — cao nhất VN'],
        ['48', 'UK School Interv 2025 (B5)', 'Editorial Q1', '—', 'CBT > mindfulness; MHST + PLACES self-referral; co-design'],
        ['49', 'JAMA App CBT 2024 (B3)', 'RCT', 'TBD', 'Mobile CBT app cho lo âu — bằng chứng RCT'],
        ['50', 'Resilience School MA 2025 (B6)', 'SR + MA RCTs', 'Nhiều RCT', 'Resilience trường có hiệu quả nhỏ-TB; heterogeneity cao'],
        ['51', 'Japan iCBT SAD 2024 (B11)', 'RCT đa trung tâm', 'VTN Nhật', 'iCBT cho subthreshold SAD — dương tính'],
        ['52', 'Academic Stress SR 2025 (A7)', 'SLR Q1 34p', 'Toàn cầu', 'CBT hiệu quả nhất; quản lý thời gian phù hợp VN'],
        ['53', 'Dinh 2021 VN School Factors (A12)', 'Cắt ngang VN', 'HS THCS', 'Yếu tố trường ảnh hưởng lo âu HS THCS VN'],
        ['54', 'Dong 2025 PLOS DASS (A5)', 'Cắt ngang TQ', '2.716 HS Ya An', 'Lo âu 41,4%; tâm sự gia đình OR=0,22 (giảm 78%)'],
        ['55*', 'Chen 2025 COVID Meta (A4)', 'Meta 3-cấp', '141 NC, 1M+ trẻ', '14 bảo vệ + 29 nguy cơ; emotional functioning bảo vệ mạnh'],
        ['56*', 'Zhang 2026 Bayesian CBT (B1)', 'Bayesian MA', '31 RCT, 19.865', 'CBT trường phổ quát hiệu quả NHỎ; chất lượng nền THẤP'],
        ['57*', 'Qiaochu 2025 Mobile CBT (B4)', 'SR 9 RCT', '2.479 trẻ', 'Trầm cảm 7/8 NC; lo âu chỉ 2/6 NC hiệu quả'],
        ['58*', 'Menon 2025 LMIC SEA (B10)', 'Scoping Review', '69 NC / 12 nước', 'KHOẢNG TRỐNG: cộng đồng + gia đình + dài hạn'],
    ],
    widths=[1.0, 4.0, 2.5, 2.5, 5.5]
)
P('* Dấu (*) = bài abstract-only (paywall); chưa có PDF đầy đủ.', italic=True, size=10)

# ============================================================
H('13. PHÁT HIỆN MỚI — Đối chiếu liên bài 36-58', level=2)

H('13.1. Bản đồ tỷ lệ lo âu VN — CẬP NHẬT', level=3)
P('Sau khi bổ sung 5 bài VN mới, bản đồ tỷ lệ lo âu HS VN đã đầy đủ hơn:', bold=True)
table(
    ['Địa bàn', 'NC', 'Mẫu', 'Công cụ', 'Tỷ lệ lo âu'],
    [
        ['Hà Nội THPT', 'Hoa 2024 (#02)', '~1.000', 'GAD-7', '40,6%'],
        ['Huế THCS (dọc 3 năm)', 'Trần Thảo Vi 2024 (#36)', '745', 'DASS', '~30-40%'],
        ['Hải Phòng THPT', '#45', '420', 'DASS-21', '39,3%'],
        ['Long An THPT', '#47', 'TBD', 'TBD', '57,2% (cao nhất)'],
        ['Vĩnh Long THPT', '#40', '919', 'CES-D + ESSA', '12,2% (chỉ trầm cảm)'],
        ['VN COVID toàn quốc SV', 'Nguyen LX 2023 (#39)', '5.730', 'GAD-7 ≥10', '16,2%'],
        ['HS THCS VN', 'Dinh 2021 (#53)', 'TBD', 'TBD', 'TBD'],
    ],
    widths=[3.5, 3.5, 2.5, 2.8, 3.2]
)
P('NHẬN XÉT: Tỷ lệ lo âu HS VN dao động 16% (SV đại học GAD-7 ≥10) → 57,2% (Long An). Sự dao động phản ánh: (1) khác biệt công cụ và cut-off, (2) khác biệt vùng (Long An ĐB SCL > Hà Nội), (3) khác biệt độ tuổi (THPT > SV). Cần NC chuẩn hóa công cụ.', italic=True)

H('13.2. CƠ CHẾ BẢO VỆ — Phát hiện mới (Dong A5 + Chen A4)', level=3)
P('Phát hiện QUAN TRỌNG NHẤT của phiên này: KÊNH GIAO TIẾP GIA ĐÌNH là yếu tố bảo vệ mạnh nhất.', bold=True)
table(
    ['Yếu tố bảo vệ', 'Hiệu lực', 'Nguồn', 'Ghi chú'],
    [
        ['Tâm sự với gia đình (vs không có)', 'OR=0,22 (giảm 78% trầm cảm)', 'Dong A5 (#54) n=2.716 TQ', 'MẠNH NHẤT — kênh giao tiếp > mức hỗ trợ chung'],
        ['Tâm sự với bạn bè', 'OR=0,37 (giảm 63%)', 'Dong A5 (#54)', 'Mạnh thứ 2'],
        ['Hỗ trợ gia đình "More" (vs hiếm)', 'OR=0,45 (giảm 55%)', 'Dong A5 (#54)', 'Mạnh thứ 3'],
        ['Tìm trợ giúp chủ động', 'OR=0,48', 'Dong A5 (#54)', 'Hành vi chủ động'],
        ['Emotional functioning', 'Yếu tố bảo vệ MẠNH', 'Chen A4 (#55) meta 141 NC', 'Chưa có hệ số cụ thể'],
        ['Family support + community resources', 'Buffer quan trọng', 'Chen A4 (#55)', 'Confirmation đa NC'],
        ['Resilience + lạc quan', 'Tăng nguồn lực tâm lý', 'B6 (#50) MA RCT + VN21', 'Effect TB; phù hợp giai đoạn 2 đề cương'],
        ['Self-efficacy (mediation)', 'MXH gián tiếp 63% qua', 'Zheng A6 (#46)', 'Cơ chế: MXH→giảm SE→tăng lo âu'],
    ],
    widths=[5.0, 4.0, 4.5, 4.0]
)
P('GỢI Ý CAN THIỆP MỚI: Tập huấn KỸ NĂNG GIAO TIẾP cha-con (chứ không chỉ "tăng hỗ trợ" chung) — phù hợp với khoảng trống Menon B10 về can thiệp gia đình ở LMIC.', bold=True)

H('13.3. CAN THIỆP CBT — Bằng chứng MÂU THUẪN', level=3)
P('Phiên này bổ sung 6 bài can thiệp (B1, B3, B4, B5, B6, B10, B11) — dẫn đến bức tranh phức tạp về CBT trường:', bold=True)
table(
    ['NC', 'Loại CBT', 'Kết quả', 'Diễn giải'],
    [
        ['Zhang B1 (#56)', 'CBT phổ quát trường', 'Hiệu quả NHỎ; chất lượng nền THẤP', 'Universal yếu — dilution'],
        ['UK B5 (#48)', 'Mindfulness trường (universal)', 'Mindfulness 8.376 HS THẤT BẠI', 'Engagement thấp; co-design cần thiết'],
        ['UK B5 (#48)', 'BESST CBT self-referral', 'TÍCH CỰC', 'Targeted + tự chọn = tốt hơn universal'],
        ['B8 Sri Lanka', 'CBT do GV (cluster RCT)', 'Lo âu giảm sau 3 tháng', 'GV có thể, dù chuyên gia lâm sàng > GV'],
        ['B6 Resilience (#50)', 'Resilience MA', 'Hiệu quả nhỏ-TB', 'Yếu tố bảo vệ — bổ sung CBT'],
        ['B11 Japan iCBT (#51)', 'iCBT subthreshold SAD', 'DƯƠNG TÍNH RCT', 'Targeted thành công'],
        ['B3 JAMA App (#49)', 'Mobile App CBT', 'RCT — chi tiết PDF', 'Bằng chứng mới'],
        ['B4 Qiaochu (#57)', 'Mobile CBT SR 9 RCT', 'Trầm cảm 7/8 NC; LO ÂU chỉ 2/6', 'Mobile mạnh trầm cảm, YẾU lo âu'],
        ['B7 CA-CBT ĐNA', 'CBT thích ứng VH', 'Có hiệu quả', 'Phù hợp LMIC'],
    ],
    widths=[3.5, 4.0, 4.5, 5.0]
)
P('TỔNG KẾT MỚI VỀ CBT TRƯỜNG:', bold=True)
P('• UNIVERSAL (phổ quát): hiệu quả NHỎ — dilution + engagement thấp (B1 Bayesian, B5 Mindfulness UK).')
P('• TARGETED + tự chọn (self-referral): TỐT HƠN universal (B5 BESST, B11 Japan iCBT).')
P('• MOBILE CBT: mạnh cho TRẦM CẢM (B4: 7/8), yếu cho LO ÂU (B4: 2/6) — nhưng SAD-specific tốt (B11 Japan, B2 JMIR g=0,878).')
P('• GV cung cấp CBT: KHẢ THI ở LMIC nhưng yếu hơn chuyên gia (B5 UK, B8 Sri Lanka).')
P('• KẾT HỢP CBT + RESILIENCE + skill thông gia đình → có thể HIỆU QUẢ HƠN single-modality.')

H('13.4. KHOẢNG TRỐNG NC — CẬP NHẬT 10/04/2026', level=3)
table(
    ['#', 'Khoảng trống MỚI', 'Bằng chứng từ', 'Ưu tiên'],
    [
        ['11', 'RCT can thiệp KỸ NĂNG GIAO TIẾP gia đình ở VN', 'Dong A5 OR=0,22 + Menon B10', 'RẤT CAO'],
        ['12', 'CBT TARGETED self-referral cho VTN VN có triệu chứng', 'B5 BESST UK + B11 Japan', 'CAO'],
        ['13', 'Mobile CBT tiếng Việt cho TRẦM CẢM (mạnh) + SAD (đặc thù)', 'B3 JAMA + B4 Qiaochu + B11', 'CAO'],
        ['14', 'NC dọc đo CẢ 3 trục DASS-21 ở mẫu VN >2.000 THPT', 'Dong A5 + VN21', 'CAO'],
        ['15', 'Cộng đồng + gia đình can thiệp ở LMIC Đông Á (gồm VN)', 'Menon B10 scoping', 'TB'],
        ['16', 'So sánh tỷ lệ lo âu vùng VN (chuẩn hóa công cụ)', '5 NC VN với tỷ lệ 16-57%', 'TB'],
        ['17', 'Resilience trường tích hợp CBT — RCT VN', 'B6 + Ireland QT32 + VN21 lạc quan', 'TB'],
    ],
    widths=[1.0, 6.5, 5.5, 2.5]
)

H('13.5. ĐỀ XUẤT BỔ SUNG ĐỀ CƯƠNG VN', level=3)
P('Dựa trên 14 bài mới, các bổ sung được đề xuất cho đề cương 3 giai đoạn:', bold=True)
P('GIAI ĐOẠN 1 — Khảo sát: Thêm câu hỏi về KÊNH GIAO TIẾP gia đình (frequency, content, confidant — như Dong A5). Đo CẢ 3 trục DASS-21. So sánh GAD-7 + DASS-21 + DASS-Y ở cùng mẫu (đã có).')
P('GIAI ĐOẠN 2 — Can thiệp: Thiết kế can thiệp PHỐI HỢP: (a) tập huấn kỹ năng giao tiếp cho cha mẹ + HS; (b) CBT TARGETED cho HS có triệu chứng (không universal); (c) module mobile CBT tiếng Việt cho trầm cảm; (d) thành phần resilience. Sử dụng thiết kế PLACES self-referral để giảm kỳ thị.')
P('GIAI ĐOẠN 3 — Phỏng vấn sâu: Tập trung tìm hiểu RÀO CẢN giao tiếp gia đình + cơ chế "fear of letting down others" (Dong A5 phát hiện 60,3% TQ).')

P('Bản cập nhật này là PHỤ LỤC bổ sung cho tổng hợp 09/04/2026. Toàn bộ 11 sections + 12 bảng cũ vẫn hợp lệ. Bài 36-44 đã được tích hợp trong sections 10-11 phiên trước.', italic=True)

doc.save(DST)
print(f'Saved: {DST}')

# Verify
d2 = Document(DST)
print(f'Total paragraphs: {len(d2.paragraphs)}, tables: {len(d2.tables)}')
