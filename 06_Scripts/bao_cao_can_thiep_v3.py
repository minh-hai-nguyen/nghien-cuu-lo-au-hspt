# -*- coding: utf-8 -*-
"""
BAO CAO CAN THIEP TAM LY RLLA VTN — v3 (11/04/2026).
Cap nhat so voi v2 (10/04/2026):
- Them 8 bai moi tim tu web + 2 bai da dich (b59, b60)
- Sua 4 loi QA (VN25, QT44, QT45, QT46)
- Phat hien QUAN TRONG: Happy House RCT Vietnam (Tran et al. 2023) — RCT VN dau tien cho VTN
- Cac phan cap nhat duoc TO DO
"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
CHARTS = os.path.join(os.path.dirname(__file__), 'charts')
OUT = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 11042026 v3.docx')
PAGE_W = 16.0

RED = RGBColor(0xCC, 0, 0)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2)

def H(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def Hred(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RED

def P(text, bold=False, italic=False, size=12, color=None, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb = color

def Pred(text, bold=False, italic=False, size=12, align='justify'):
    """Paragraph with RED color — for updates vs v2"""
    return P(text, bold=bold, italic=italic, size=size, color=RED, align=align)

def table(headers, rows, widths, red=False):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
                if red: r.font.color.rgb = RED
        shade(c, 'FFE4E1' if red else 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
                    if red: r.font.color.rgb = RED
    return t

def img(filename, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(CHARTS, filename)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.italic=True

# ============================================================
# TRANG TIÊU ĐỀ
# ============================================================
H('BÁO CÁO TỔNG HỢP', level=1)
H('BIỆN PHÁP HỖ TRỢ TÂM LÝ VÀ CAN THIỆP TÂM LÝ\nCHO TRẺ VỊ THÀNH NIÊN BỊ RỐI LOẠN LO ÂU', level=1)
P('Phiên bản 3 (cập nhật 11/04/2026)', italic=True, align='center')
Pred('NHỮNG CẬP NHẬT QUAN TRỌNG SO VỚI V2 (10/04/2026) ĐƯỢC TÔ MÀU ĐỎ', bold=True, italic=True, align='center')
P('Tổng hợp 23 công trình nghiên cứu — Phân theo 3 vùng địa lý:\n'
  '(1) Việt Nam | (2) Châu Á (ngoài Việt Nam) | (3) Châu Âu — Châu Úc — Châu Mỹ',
  italic=True, align='center')

# ============================================================
# GHI CHÚ CẬP NHẬT V3
# ============================================================
Hred('NHỮNG CẬP NHẬT TRONG V3 (so với v2 ngày 10/04/2026)', level=2)
Pred('Bản v3 bổ sung và sửa đổi như sau so với v2:', bold=True)
Pred('1. PHÁT HIỆN QUAN TRỌNG NHẤT — Happy House RCT Việt Nam (Tran, Nguyen, Shochet et al. 2023, '
     'Cambridge Prisms: Global Mental Health): đây là RCT trường học ĐẦU TIÊN cho VTN tại Việt '
     'Nam (n = 1.084 HS lớp 10 Hà Nội) — TRÁI với khẳng định trong v2 rằng "0 RCT VN cho VTN". '
     'Cần cập nhật mạnh phần Việt Nam.')
Pred('2. Sửa 4 lỗi tìm thấy qua kiểm tra QA: (a) VN25 Hải Phòng tỷ lệ lo âu thực tế là 53,81% '
     '(THPT Cộng Hiền) hoặc 43,33% (GDNN-GDTX Vĩnh Bảo), không phải 39,3% như v2; (b) QT44 tác giả '
     'là Cai (không phải Cao); (c) QT45 Sasaki ID đăng ký UMIN000050064 (không phải UMIN000049768); '
     '(d) QT46 bổ sung tác giả Jagiello et al. 2025.')
Pred('3. Bổ sung 8 nghiên cứu mới tìm từ web 2024–2025: Happy House VN, Tran Duc Si Long An 2025, '
     'School Nurse SR (tandfonline 2025), CoolMinds ICBT 2025, ICBT stepped care videoconferencing '
     '(JMIR Mental Health 2025), MIXCS Japan CA-CBT 2025, VR exposure therapy SR+MA 2025, '
     'Physical activity umbrella review 2025 (JAACAP).')
Pred('4. Bổ sung 2 bản dịch mới (bài 59, 60): Đào Thị Ngoãn 2025 TCNCYH (SV năm 4 ĐH Y Hà Nội) + '
     'Duong et al. 2025 (Social Psychiatry Q1, n = 2.631 HS TPHCM).')
Pred('5. Cập nhật bảng xếp hạng can thiệp + đề xuất 5 thành phần can thiệp — mở rộng thêm cấp độ '
     'phụ huynh (parent-child communication) dựa vào Zapf 2024 JCPP Advances + Menon 2025.')

P('Các phần KHÔNG được tô đỏ giữ nguyên từ v2 — để tiện so sánh.', italic=True, size=11)

# ============================================================
# TÓM LƯỢC ĐIỀU HÀNH (sửa 1 điểm — xoá "0 RCT" + thêm Happy House)
# ============================================================
H('TÓM LƯỢC ĐIỀU HÀNH', level=2)

P('Báo cáo này tổng hợp 23 công trình nghiên cứu về biện pháp hỗ trợ và can thiệp tâm lý '
  'cho rối loạn lo âu (RLLA) ở trẻ vị thành niên (VTN), được phân theo ba vùng địa lý nhằm '
  'làm rõ bằng chứng khu vực, mức độ phù hợp văn hoá và khả năng áp dụng cho Việt Nam. '
  'Phương pháp tổng hợp kết hợp đối chiếu liên bài, phân tích phản biện và xếp hạng bằng chứng '
  'theo chất lượng thiết kế nghiên cứu.')

P('Sáu phát hiện then chốt:', bold=True)
P('Thứ nhất, Liệu pháp Nhận thức – Hành vi (CBT) là phương pháp được chứng minh hiệu quả nhất '
  'và nhất quán xuyên suốt cả ba vùng địa lý. Bằng chứng vàng từ thử nghiệm CAMS '
  '(Walkup et al. 2008, NEJM 359:2753–66) cho thấy CBT kết hợp Sertraline đạt 80,7 % đáp ứng '
  'so với 59,7 % cho CBT đơn thuần và 23,7 % cho giả dược (n = 488 trẻ 7–17 tuổi).')
P('Thứ hai, CBT qua Internet (iCBT) cho rối loạn lo âu xã hội (SAD) đứng hạng 1 trong phân tích '
  'tổng hợp mạng 30 RCT (Xian, Zhang & Jiang 2024, J Affect Disord 365:614–627; SUCRA 71,2 %), '
  'và kích thước hiệu ứng đặc biệt LỚN khi can thiệp được thiết kế riêng cho SAD '
  '(Walder et al. 2025, JMIR; Hedges g = 0,878).')
P('Thứ ba, ứng dụng di động CBT (Mobile CBT) cho hiệu quả MẠNH với trầm cảm (87,5 % nghiên cứu '
  'dương tính) nhưng YẾU với lo âu khi không thiết kế riêng (Qiaochu & Wang 2025, Clin Psychol '
  'Psychother 32:e70173). Tuy nhiên RCT ứng dụng Maya của Bress et al. 2024 cho thấy hiệu lực '
  'rất lớn (Cohen d = 1,04 sau 12 tuần, JAMA Network Open 7(8):e2428372).')
P('Thứ tư, can thiệp PHỔ QUÁT tại trường (universal CBT) chỉ cho hiệu quả nhỏ và chất lượng '
  'bằng chứng nền thấp (Zhang, Liang & Kang 2026, J Clin Psychol 82:248–259, Bayesian MA 31 RCT, '
  'n = 19.865). Mô hình CHỈ ĐỊNH (targeted) cho học sinh có triệu chứng — như BESST UK '
  '(Brown & Carter 2025, J Mental Health 34(4):357–361) — hứa hẹn hơn và phù hợp với cảnh báo '
  'của Brown & Carter về thất bại Mindfulness phổ quát ở 8.376 học sinh UK.')

Pred('Thứ năm (CẬP NHẬT V3): TẠI VIỆT NAM đã có một RCT trường học cho VTN — thử nghiệm '
     '"Happy House" của Tran, Nguyen, Shochet và cộng sự (2023, Cambridge Prisms: Global '
     'Mental Health, xuất bản 10/2023). Đây là chương trình Resourceful Adolescent Program '
     '(RAP-A) được thích ứng văn hoá cho HS Việt Nam, 6 buổi mỗi buổi 90 phút, n = 1.084 HS lớp '
     '10 Hà Nội (8 trường: 4 can thiệp + 4 đối chứng). Hiệu ứng trên trầm cảm: Cohen d = 0,11 '
     '(p = 0,011) ở 2 tuần sau can thiệp; OR = 0,56 (p = 0,027) cho trầm cảm có ý nghĩa lâm '
     'sàng. Tuy nhiên, tác dụng KHÔNG CÒN ý nghĩa thống kê ở 6 tháng theo dõi (có thể do ảnh '
     'hưởng COVID-19). KHOẢNG TRỐNG CÒN LẠI: RCT CBT TARGETED cho HS có triệu chứng (Happy '
     'House là UNIVERSAL), RCT cho rối loạn lo âu chuyên biệt, RCT cho THPT lớp 11–12 và THCS.')

P('Thứ sáu, khu vực Đông Á và Thái Bình Dương ở các nước có thu nhập trung bình – thấp (LMIC) '
  'còn thiếu nghiêm trọng các can thiệp ở cấp CỘNG ĐỒNG, GIA ĐÌNH và dịch vụ ĐÁP ỨNG dài hạn '
  '(Menon et al. 2025, Asia Pac J Public Health 37(4):332–346, scoping review 69 nghiên cứu, '
  '12 quốc gia).')

# ============================================================
# PHẦN I — VIỆT NAM (cập nhật mạnh)
# ============================================================
doc.add_page_break()
H('PHẦN I — VIỆT NAM', level=1)

Pred('CẬP NHẬT V3: Trong v2, phần này khẳng định "0 RCT can thiệp tâm lý cho VTN VN" — KHẲNG '
     'ĐỊNH NÀY KHÔNG CHÍNH XÁC. Sau khi tìm kiếm sâu trên web, chúng tôi phát hiện thử nghiệm '
     'Happy House (Tran et al. 2023, Cambridge Prisms: Global Mental Health) là RCT trường học '
     'ĐẦU TIÊN cho VTN tại Việt Nam. Phần I được viết lại để bao gồm phát hiện này bên cạnh '
     'luận án Trần Nguyễn Ngọc 2018 (đối tượng người lớn nội trú).', italic=True)

# SECTION 1.1 — Happy House RCT (NEW in v3)
Hred('1.1. Tran, Nguyen, Shochet et al. (2023) — Happy House RCT Việt Nam (BÀI MỚI TRONG V3)', level=2)

Hred('1.1.1. Thông tin thư mục', level=3)
table(
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Thach Duc Tran, Huong Nguyen, Ian Shochet, Nga Nguyen, Nga La, Astrid Wurfl, Jayne Orr, Hau Nguyen, Ruby Stocker, Jane Fisher'],
        ['Đơn vị', 'Monash University (Úc), Queensland University of Technology (Úc), và các đơn vị đối tác Việt Nam tại Hà Nội'],
        ['Tiêu đề', 'School-based universal mental health promotion intervention for adolescents in Vietnam: Two-arm, parallel, controlled trial'],
        ['Tạp chí', 'Cambridge Prisms: Global Mental Health'],
        ['Năm', '10/2023'],
        ['Link', 'https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/'],
        ['Loại NC', 'Cluster RCT hai nhánh song song có đối chứng'],
        ['Mẫu', '1.084 HS lớp 10 (531 can thiệp + 552 đối chứng), 8 trường THPT Hà Nội (4 can thiệp + 4 đối chứng)'],
        ['Can thiệp', 'Happy House — chương trình RAP-A (Resourceful Adolescent Program) đã được THÍCH ỨNG VĂN HOÁ Việt Nam'],
        ['Thời lượng', '6 buổi, mỗi buổi 90 phút, trong 6 tuần'],
        ['Cơ chế lý thuyết', 'CBT + liệu pháp giữa cá nhân (IPT) — kỹ năng "positive self-talk", giữ bình tĩnh, giải quyết vấn đề, mạng lưới hỗ trợ'],
        ['Người cung cấp', 'Giáo viên đã được đào tạo, dạy theo nhóm lớp'],
        ['Công cụ đo', 'Thang trầm cảm (chính), thang phúc lợi tâm lý, coping self-efficacy'],
        ['Thời điểm đo', 'Trước can thiệp, 2 tuần sau, 6 tháng sau'],
    ],
    widths=[4.0, 11.5], red=True)

Hred('1.1.2. Kết quả chính', level=3)
table(
    ['Kết quả', '2 tuần sau', '6 tháng sau'],
    [
        ['Trầm cảm (Cohen d)', 'd = 0,11 (p = 0,011) *', 'không còn ý nghĩa'],
        ['Tỷ lệ trầm cảm lâm sàng', 'OR = 0,56 (p = 0,027) *', 'không còn ý nghĩa'],
        ['Phúc lợi tâm lý', 'd = 0,13 (p < 0,05) *', 'không còn ý nghĩa'],
        ['Coping self-efficacy', 'd = 0,17–0,26 *', 'd = 0,17–0,26 * DUY TRÌ'],
    ],
    widths=[5.5, 5.0, 5.0], red=True)
Pred('* có ý nghĩa thống kê. Nguồn: Tran et al. 2023, Cambridge Prisms Global Mental Health, '
     'https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/', italic=True, size=10)

Hred('1.1.3. Quan điểm phản biện', level=3)
Pred('Điểm mạnh:', bold=True)
Pred('• Đây là RCT CLUSTER đầu tiên về can thiệp SKTT học đường cho VTN Việt Nam — lấp khoảng '
     'trống nghiên cứu cực lớn mà dự án của chúng ta đã xác định trước đó.')
Pred('• Cỡ mẫu rất lớn (n = 1.084) — đủ power cho phân tích hồi quy đa biến và subgroup.')
Pred('• Chương trình RAP-A có cơ sở bằng chứng quốc tế mạnh (Úc) + được thích ứng văn hoá Việt '
     'Nam một cách hệ thống (xem Nguyen et al. 2022 PLOS ONE về quá trình adaptation).')
Pred('• Tác giả liên kết Việt–Úc (Monash, QUT + Hà Nội) — mô hình hợp tác quốc tế mẫu.')
Pred('• Coping self-efficacy DUY TRÌ ở 6 tháng — cho thấy can thiệp có hiệu quả dài hạn trên '
     'khả năng ứng phó, dù không còn giảm triệu chứng.')
Pred('• Trẻ em/VTN được can thiệp là LỚP 10 Hà Nội — đúng nhóm tuổi và đô thị trung tâm VN.')

Pred('Hạn chế:', bold=True)
Pred('• Đây là can thiệp UNIVERSAL (phổ quát toàn lớp), KHÔNG phải TARGETED cho HS có triệu '
     'chứng. Hiệu ứng nhỏ (d = 0,11) có thể do dilution — phù hợp với cảnh báo của Zhang, Liang '
     '& Kang 2026 (Bayesian MA 31 RCT universal CBT).')
Pred('• Hiệu ứng KHÔNG CÒN ý nghĩa ở 6 tháng — có thể do COVID-19 gián đoạn hoặc do liều '
     'can thiệp quá ngắn (6 buổi × 90 phút). Cần booster session.')
Pred('• Đo TRẦM CẢM chính, KHÔNG đo LO ÂU riêng — chưa có bằng chứng cho rối loạn lo âu cụ thể.')
Pred('• Chỉ Hà Nội + lớp 10 — chưa đại diện toàn VN (cần lặp lại ở THCS, THPT lớp 11–12, vùng '
     'nông thôn, DTTS).')
Pred('• Giáo viên cung cấp — chưa so sánh với chuyên gia tâm lý (khác mô hình Sri Lanka De '
     'Silva 2024 nơi GV hiệu quả).')
Pred('• KHÔNG có nhóm active control — chỉ là usual curriculum, nên khó loại trừ hiệu ứng chú ý.')

Hred('1.1.4. Áp dụng cho đề cương Việt Nam', level=3)
Pred('Happy House cung cấp bài học quý cho thiết kế đề cương tiếp theo:')
Pred('(1) Mô hình RAP-A thích ứng văn hoá VN đã KHẢ THI — tiết kiệm chi phí phát triển nội dung '
     'mới, có thể xây dựng tiếp.')
Pred('(2) Cần BOOSTER SESSIONS ở 3, 6 tháng để duy trì hiệu quả — giải quyết vấn đề "fade-out" '
     'mà Happy House gặp phải.')
Pred('(3) Nên thử nghiệm mô hình TARGETED (HS có GAD-7 ≥ 8 hoặc DASS-Y lo âu ≥ 8) thay vì universal '
     '— có thể có hiệu ứng lớn hơn.')
Pred('(4) Đo đồng thời LO ÂU (GAD-7, DASS-Y, SIAS-17) + trầm cảm + coping self-efficacy — bao phủ '
     'cả giảm triệu chứng và tăng nguồn lực bảo vệ.')
Pred('(5) Mở rộng địa bàn ra ngoài Hà Nội (TPHCM, Đà Nẵng, vùng miền trung nông thôn) để tăng '
     'tính đại diện.')

Pred('Đánh giá: ⭐⭐⭐⭐ Cao. RCT VN đầu tiên, cluster thiết kế mạnh, mẫu lớn, tác giả Monash–QUT '
     'uy tín. Hạn chế: universal interventions với hiệu ứng nhỏ, không đo lo âu riêng.', bold=True)

# SECTION 1.2 — Trần Nguyễn Ngọc 2018 (giữ nguyên từ v2)
H('1.2. Trần Nguyễn Ngọc (2018) — Liệu pháp Thư giãn – Luyện tập cho RLLA lan toả', level=2)
P('(Giữ nguyên nội dung chi tiết từ v2 — xem bản v2 cho thông tin đầy đủ.)', italic=True)
P('Tóm tắt nhanh: Luận án TSYH ĐH Y Hà Nội, PGS.TS. Nguyễn Kim Việt hướng dẫn, 170 bệnh nhân '
  'RLLA lan toả (F41.1) nội trú Viện SKTT Bệnh viện Bạch Mai. Can thiệp: Liệu pháp Thư giãn – '
  'Luyện tập 20 buổi/4 tuần. Thiết kế trước–sau (n=99 hoàn thành). Kết quả: HAM-A nặng giảm từ '
  '45,5% (T0) xuống 11,1% (T4), p < 0,0001. Đối tượng là NGƯỜI LỚN, không phải VTN — cần '
  'adaptation.')

# ============================================================
# PHẦN II — CHÂU Á (cập nhật: thêm MIXCS Japan 2025)
# ============================================================
doc.add_page_break()
H('PHẦN II — CHÂU Á (ngoài Việt Nam)', level=1)

H('2.1. Tổng quan nghiên cứu Châu Á (cập nhật v3)', level=2)

table(
    ['#', 'Tác giả / Năm', 'Quốc gia', 'Loại NC', 'Cỡ mẫu', 'Phát hiện then chốt'],
    [
        ['41', 'Praptomojati & Hartanto 2024', 'Tổng hợp ĐNA', 'SR 7 NC', 'Trẻ + VTN', 'CA-CBT thích ứng VH; 0/7 NC từ VN'],
        ['42', 'De Silva et al. 2024', 'Sri Lanka', 'Cluster RCT', '720 HS lớp 9', 'CBT do GV; β = −0,096'],
        ['43', 'Xian, Zhang & Jiang 2024', 'Trung Quốc', 'NMA Bayesian', '30 RCT, 1.547 VTN', 'iCBT SUCRA 71,2% hạng 1'],
        ['51', 'Sasaki et al. 2024', 'Nhật Bản', 'RCT đa trung tâm', '77 HS-SV', 'iCBT subthreshold SAD; OR=4,97; UMIN000050064'],
        ['57', 'Qiaochu & Wang 2025', 'Trung Quốc', 'SR 9 RCT', '2.479', 'Mobile CBT: trầm cảm 7/8; lo âu 2/6'],
        ['58', 'Menon et al. 2025', 'LMIC ĐÁ-TBD', 'Scoping 69 NC', '12 quốc gia', 'Thiếu cộng đồng/gia đình/dài hạn'],
        ['29', 'Li et al. 2025', 'Trung Quốc', 'NMA Bayesian', '30 RCT, 1.711', 'CBT SUCRA 0,66; PE 0,51'],
    ],
    widths=[0.8, 3.2, 1.8, 2.5, 2.5, 5.0])

Hred('Bài mới tìm qua web (2025): MIXCS Japan', level=3)
Pred('Trial MIXCS Japan (Multi-, Inter-, and Cross-cultural Clinical Child Study) của nhóm '
     'nghiên cứu Nhật Bản — thử nghiệm so sánh CA-CBT (Culturally Adapted) vs PA-CBT '
     '(Programme-Adopted) cho rối loạn lo âu trẻ em Nhật Bản, đa trung tâm 3 tỉnh (Kyoto, '
     'Hyogo, Nagano), tuyển từ 01/2022 đến 03/2025. Thiết kế 3 nhánh: CA-CBT, PA-CBT, và '
     'psychological control. Giao thức: PMC10357790.')
Pred('Ý nghĩa cho VN: cung cấp mô hình mới cho việc thích ứng CBT cho trẻ em châu Á. Cần '
     'theo dõi kết quả khi công bố 2025-2026.')

H('2.2–2.8. Chi tiết các bài QT37–QT51 + Li 2025 BMC NMA', level=2)
P('(Giữ nguyên nội dung chi tiết từ v2 — xem bản v2 cho thông tin đầy đủ về Praptomojati 2024, '
  'De Silva 2024, Xian 2024, Walder 2025 (thuộc phần III), Sasaki 2024, Qiaochu 2025, Menon '
  '2025, Cai 2025 Resilience, Li 2025 BMC NMA.)', italic=True)

Pred('SỬA LỖI V3:', bold=True)
Pred('• QT44 Cai Resilience (không phải Cao): tác giả chính là Chenyi Cai, Southwest University '
     'Trùng Khánh. Tên đã được sửa trong tóm tắt QT44_Cai_Resilience_MA_2025.docx.')
Pred('• QT45 Sasaki: ID đăng ký thử nghiệm đúng là UMIN000050064 (không phải 000049768). Đã sửa.')
Pred('• QT46 Academic Stress SR: bổ sung tác giả Jagiello T, Belcher J, Neelakandan A, Boyd K, '
     'Wuthrich VM (2025). Tổng 31 nghiên cứu từ 13 quốc gia (không phải "34 trang" như v2 — 34 '
     'là số trang bài báo).')

# ============================================================
# PHẦN III — ÂU – ÚC – MỸ (cập nhật: thêm CoolMinds, stepped care, VR, PA umbrella, school nurse)
# ============================================================
doc.add_page_break()
H('PHẦN III — CHÂU ÂU – CHÂU ÚC – CHÂU MỸ', level=1)

H('3.1. Tổng quan các nghiên cứu Âu–Úc–Mỹ', level=2)
P('(Các bài 28 Zugman AJP, 29 Li BMC, 44 Walder JMIR, 48 Brown & Carter UK, 49 Bress JAMA, '
  '50 Cai Frontiers, 56 Zhang Bayesian, CAMS Walkup NEJM — giữ nguyên từ v2.)', italic=True)

Hred('3.2. BÀI MỚI V3 — CoolMinds ICBT + Stepped Care videoconferencing (2025)', level=2)
Pred('Phát hiện từ tìm kiếm web: 2 nghiên cứu mới về iCBT 2025 bổ sung quan trọng cho '
     'Phần III:', bold=True)

Pred('(A) CoolMinds ICBT — Đan Mạch (Studsgaard et al. 2025)', bold=True)
Pred('• Tiêu đề: "Developing an Internet-Based Cognitive Behavioral Therapy Intervention for '
     'Adolescents With Anxiety Disorders: Design, Usability, and Initial Evaluation of the '
     'CoolMinds Intervention"')
Pred('• Tạp chí: JMIR Formative Research (2025, vol. 9, e66966)')
Pred('• Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC12015348/')
Pred('• Đối tượng: VTN 12-17 tuổi có rối loạn lo âu lâm sàng, Đan Mạch')
Pred('• Phương pháp: User-centered design 3 pha (identify needs → prototype → feasibility). '
     'ICBT dành cho VTN, dùng giao diện trực quan phù hợp lứa tuổi.')
Pred('• Ý nghĩa: mô hình phát triển app có người dùng cùng thiết kế (co-design) — phù hợp '
     'nguyên tắc Brown & Carter 2025 về co-design. Bước tiếp theo là RCT feasibility '
     '(Lautrup et al. 2025, Scand J Cogn Behav Ther 2025).')

Pred('(B) Stepped Care ICBT với Videoconferencing — Úc (Baumgart et al. 2025)', bold=True)
Pred('• Tiêu đề: "Integrating Videoconferencing Therapist Guidance Into Stepped Care '
     'Internet-Delivered Cognitive Behavioral Therapy for Child and Adolescent Anxiety: '
     'Noninferiority Randomized Controlled Trial"')
Pred('• Tạp chí: JMIR Mental Health (2025, vol. 12, e57405)')
Pred('• Link: https://mental.jmir.org/2025/1/e57405')
Pred('• Mẫu: 137 trẻ + VTN Úc 7-17 tuổi có rối loạn lo âu chính (primary anxiety disorder)')
Pred('• Thiết kế: RCT noninferiority so sánh (1) ICBT-SC[VC] — stepped care với videoconferencing '
     'cho step-up vs (2) ICBT-TG[VC] — full therapist delivery qua videoconference.')
Pred('• Kết quả: Stepped care đạt remission rate TƯƠNG ĐƯƠNG full therapist — noninferiority '
     'xác nhận. Chi phí stepped care THẤP HƠN vì chỉ step-up khi cần.')
Pred('• Ý nghĩa cho VN: mô hình stepped care có thể là giải pháp cho các cơ sở y tế thiếu '
     'chuyên gia tâm lý — bắt đầu với iCBT tự học, chỉ tăng bậc khi cần có therapist.')

Hred('3.3. BÀI MỚI V3 — School Nurse-Delivered Anxiety SR+MA 2025', level=2)
Pred('Tiêu đề: "School nurse–delivered anxiety interventions for adolescents: a systematic '
     'review and meta-analysis"')
Pred('Tác giả: nhóm NC tandfonline, xuất bản 18/06/2025')
Pred('Tạp chí: Cogent Psychology (hoặc tương đương Taylor & Francis)')
Pred('Link: https://www.tandfonline.com/doi/full/10.1080/23311908.2025.2519529')
Pred('Phương pháp: SR + MA theo PRISMA 2020 + protocol đăng ký. Đối tượng VTN 12-18 tuổi. '
     'Can thiệp do y tá học đường cung cấp cho lo âu.')
Pred('Kết quả: Cohen d = 0,65 cho youth-rated anxiety outcomes (post-treatment). Kích thước '
     'hiệu ứng TRUNG BÌNH-LỚN. Bổ sung bằng chứng Imondi et al. 2025 (Child Anxiety Learning '
     'Modules — CALM — do y tá học đường cung cấp, đối tượng tiểu học, NASN School Nurse '
     'vol. 40, 10.1177/10598405251328369).')
Pred('Ý nghĩa cho VN: mô hình Y TÁ HỌC ĐƯỜNG là CẦU NỐI giữa mô hình Sri Lanka (GV) và mô '
     'hình UK (chuyên gia tâm lý). Việt Nam có hệ thống y tá học đường — có thể phát triển '
     'đào tạo đặc biệt cho nhóm này.')

Hred('3.4. BÀI MỚI V3 — VR Exposure Therapy SR+MA 2025', level=2)
Pred('Tiêu đề: "Effectiveness of virtual reality therapy in the treatment of anxiety disorders '
     'in adolescents and adults: a systematic review and meta-analysis of randomized controlled '
     'trials"')
Pred('Tạp chí: Frontiers in Psychiatry (2025, vol. 16, article 1553290)')
Pred('Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC11904249/')
Pred('Mẫu: 33 RCT, n = 3.182 VTN và người lớn có rối loạn lo âu.')
Pred('Kết quả: VR therapy CẢI THIỆN ĐÁNG KỂ triệu chứng lo âu so với can thiệp thông thường. '
     'Đặc biệt hiệu quả cho SPECIFIC PHOBIAS (ám ảnh cụ thể) và SOCIAL ANXIETY. Hiệu lực TƯƠNG '
     'ĐƯƠNG với can thiệp truyền thống (face-to-face exposure).')
Pred('Bổ sung (VIRTUS trial 2025 BMC Psychiatry): VR exposure training cho VTN có lo âu xã hội '
     '— RCT, giao thức đã công bố. Pilot study (Zimmer et al. 2024, JMIR Mental Health 11:e56235) '
     'cho thấy VRET giảm state anxiety + social anxiety sau can thiệp.')
Pred('Ý nghĩa cho VN: VR còn tốn kém nhưng đang rẻ dần. Mô hình có tiềm năng cho can thiệp '
     'phơi nhiễm (exposure therapy) — đặc biệt cho phobias chuyên biệt và SAD. Có thể thử '
     'nghiệm ở các thành phố lớn.')

Hred('3.5. BÀI MỚI V3 — Physical Activity Umbrella Review 2025 (JAACAP)', level=2)
Pred('Tiêu đề: "Systematic Umbrella Review and Meta-Meta-Analysis: Effectiveness of Physical '
     'Activity in Improving Depression and Anxiety in Children and Adolescents"')
Pred('Tạp chí: Journal of the American Academy of Child & Adolescent Psychiatry (JAACAP, '
     '2025 — bằng chứng TỐT NHẤT về PE cho VTN đến nay)')
Pred('Link: https://www.jaacap.org/article/S0890-8567(25)00208-4/fulltext')
Pred('Mẫu: UMBRELLA REVIEW 21 systematic reviews, tổng 375 RCTs, n = 38.117 người tham gia '
     '(5–18 tuổi). Đây là tổng hợp LỚN NHẤT về physical activity cho SKTT trẻ em/VTN.')
Pred('Kết quả: SMD = −0,39 cho LO ÂU (hiệu ứng trung bình). Resistance exercise hiệu quả nhất '
     'cho lo âu. Aerobic exercise SMD = −0,32 cho lo âu. Interventions < 12 tuần hiệu quả hơn '
     'cho trầm cảm.')
Pred('Ý nghĩa cho VN: BẰNG CHỨNG VÀNG cho PE — hiệu ứng trung bình và lớn hơn nhiều bài khác. '
     'Tích hợp PE vào chương trình thể dục học đường có sẵn. Ưu tiên resistance + aerobic. '
     'Thời lượng < 12 tuần phù hợp semester học.')

# ============================================================
# PHẦN IV — TỔNG KẾT (cập nhật bảng xếp hạng + 5 thành phần)
# ============================================================
doc.add_page_break()
H('PHẦN IV — TỔNG KẾT VÀ KHUYẾN NGHỊ CHO VIỆT NAM', level=1)

H('4.1. So sánh 3 vùng địa lý (cập nhật v3)', level=2)
table(
    ['Tiêu chí', 'Việt Nam', 'Châu Á (ngoài VN)', 'Châu Âu – Úc – Mỹ'],
    [
        ['Số bài can thiệp', '2 *', '8 (thêm MIXCS)', '11 (thêm 4 bài mới v3)'],
        ['Đối tượng VTN', 'CÓ (Happy House) + KHÔNG (TNNgoc)', 'CÓ', 'CÓ'],
        ['Loại NC chủ yếu', 'RCT cluster (Happy House) + trước–sau', 'RCT + NMA + cluster RCT', 'RCT + MA + umbrella review'],
        ['Phương pháp chính', 'RAP-A thích ứng + Thư giãn-Luyện tập', 'CBT + iCBT + Mobile + CA-CBT', 'CBT + iCBT + DMHI + VR + PE + school nurse'],
        ['Bối cảnh', 'Trường THPT Hà Nội + BV Bạch Mai', 'Trường + BV + Online', 'Trường + Online + Phòng khám'],
        ['Bằng chứng cao nhất', 'Trung bình (Happy House d=0,11)', 'Mạnh (Xian NMA)', 'Rất mạnh (CAMS NEJM, umbrella review JAACAP 2025)'],
        ['Khoảng trống', 'CÒN LỚN — thiếu TARGETED + LO ÂU chuyên biệt', 'Vừa', 'Nhỏ'],
    ],
    widths=[3.5, 4.0, 4.0, 4.0])
Pred('* Happy House (Tran et al. 2023) + Trần Nguyễn Ngọc 2018 (người lớn). Trong v2 đếm 1 — '
     'thiếu Happy House.', italic=True, size=10)

H('4.2. Bảng xếp hạng 15 phương pháp can thiệp (cập nhật v3)', level=2)
table(
    ['Hạng', 'Phương pháp', 'Bằng chứng tốt nhất', 'Phù hợp VN', 'Khuyến nghị'],
    [
        ['1', 'CBT + SSRI kết hợp', 'CAMS NEJM 80,7 %', 'Hạn chế chi phí', 'Chỉ ca nặng'],
        ['2', 'CBT cá nhân / nhóm', 'CAMS 59,7 %; Li 2025 SUCRA 0,66', 'CỐT LÕI', 'ĐẦU TƯ đào tạo'],
        ['3', 'iCBT cho SAD', 'Xian NMA hạng 1; Walder g = 0,878', 'CAO — số hoá', 'Phát triển app VN'],
        ['4', 'gCBT (CBT nhóm)', 'Xian NMA hạng 2 chức năng', 'CAO', 'Triển khai trường'],
        ['5', 'CBT do GV (school-based)', 'De Silva 2024 RCT 720 HS; Happy House VN', 'CỐT LÕI LMIC', 'Đào tạo GV + booster'],
        ['6', 'CA-CBT (thích ứng VH)', 'Praptomojati 2024 SR; MIXCS Japan 2025', 'CỰC CAO', 'Phát triển bản VN'],
        ['7', 'PE (hoạt động thể chất)', 'JAACAP 2025 umbrella SMD = −0,39', 'CAO — đã sẵn', 'Giờ thể dục + resistance'],
        ['8', 'Mobile App CBT (Maya-style)', 'Bress 2024 d = 1,04', 'CAO', 'RCT pilot'],
        ['9', 'Resilience training', 'Cai 2025 MA Frontiers', 'CAO', 'Module bổ trợ'],
        ['10', 'Thư giãn – Luyện tập (Yoga)', 'Trần Nguyễn Ngọc 2018', 'CAO — đã thử VN', 'Module bổ trợ'],
        ['11', 'Mobile CBT (trầm cảm)', 'Qiaochu 2025 87,5 %', 'CAO trầm cảm', 'Bắt đầu với trầm cảm'],
        ['12', 'VR exposure therapy (VRET)', 'Frontiers 2025 MA 33 RCT SMD đáng kể', 'TB — cần công nghệ', 'Thử nghiệm đô thị'],
        ['13', 'School nurse–delivered', 'Tandfonline SR+MA 2025 d = 0,65', 'CAO — hệ thống có sẵn', 'Đào tạo y tá học đường'],
        ['14', 'PLACES self-referral', 'Brown & Carter 2025 BESST', 'TB', 'Thử pilot'],
        ['15', 'Mindfulness phổ quát', 'Brown 2025: 8.376 HS thất bại; Dunning re-analysis 2024', 'KHÔNG', 'TRÁNH universal'],
    ],
    widths=[1.0, 4.0, 4.5, 3.0, 3.0])
Pred('CẬP NHẬT V3: Thêm hạng 12 (VRET), 13 (school nurse), cập nhật hạng 5 với Happy House VN + '
     'booster, hạng 7 với umbrella review JAACAP, hạng 9 Cai (không Cao).', italic=True, size=10)

H('4.3. Đề xuất thiết kế can thiệp 5 thành phần (cập nhật v3)', level=2)
Pred('CẬP NHẬT: bổ sung "booster sessions" (bài học từ Happy House fade-out ở 6 tháng) và '
     'mở rộng module giao tiếp gia đình với cấp độ cha mẹ.', italic=True)
table(
    ['Thành phần', 'Số buổi', 'Cơ sở bằng chứng (v3)', 'Áp dụng VN'],
    [
        ['1. CBT NHÓM TARGETED', '8 buổi × 60 phút', 'CAMS 2008; Li 2025; De Silva 2024; Happy House VN 2023', 'Cốt lõi; TARGETED (khác Happy House universal)'],
        ['2. Module GIAO TIẾP GIA ĐÌNH', '4 buổi cha mẹ + HS', 'Dong 2025 (OR=0,22); Menon 2025; Zapf 2024', 'Module MỚI: kỹ năng "tâm sự"'],
        ['3. Module RESILIENCE', '3 buổi', 'Cai 2025; Trần Thảo Vi 2025 (lạc quan)', 'Lạc quan, kết nối XH'],
        ['4. PE + Thư giãn – Thở', '8 buổi tích hợp thể dục', 'JAACAP umbrella 2025 (SMD=−0,39); Trần Nguyễn Ngọc 2018', 'Resistance + aerobic < 12 tuần'],
        ['5. App iCBT tiếng Việt', 'Suốt 12 tuần + videoconf hỗ trợ', 'Walder 2025 g=0,878; Sasaki 2024; CoolMinds 2025; Baumgart 2025 stepped care', 'Stepped care: iCBT tự học → video call khi cần'],
        ['BOOSTER', '2 buổi ở 3 và 6 tháng', 'Happy House 2023 fade-out', 'Tránh mất hiệu ứng dài hạn'],
    ],
    widths=[3.5, 2.8, 6.5, 2.7], red=True)

H('4.4. Năm khoảng trống nghiên cứu tại Việt Nam (cập nhật v3)', level=2)
Pred('CẬP NHẬT — sau khi phát hiện Happy House, khoảng trống số 1 đã được LẤP MỘT PHẦN '
     '(universal), nhưng xuất hiện khoảng trống mới:')
Pred('1. [CẬP NHẬT] Thay vì "0 RCT can thiệp" → còn thiếu RCT TARGETED cho HS có triệu chứng '
     'lo âu rõ rệt (Happy House là universal trên lớp 10, không phải targeted). Cần RCT cho '
     'nhóm GAD-7 ≥ 8 hoặc DASS-Y lo âu ≥ 8.')
Pred('2. Phát triển app iCBT / Mobile CBT TIẾNG VIỆT — phù hợp xếp hạng số 3, 4, 8, 11 + '
     'mô hình stepped care (Baumgart 2025).')
Pred('3. Đào tạo GIÁO VIÊN HOẶC Y TÁ HỌC ĐƯỜNG cung cấp CBT — mô hình De Silva 2024 (Sri Lanka), '
     'Happy House VN 2023 (đã có), và tandfonline 2025 SR+MA school nurse.')
Pred('4. CA-CBT (CBT thích ứng văn hoá Á Đông) phiên bản VN CHUYÊN BIỆT CHO LO ÂU — Happy '
     'House RAP-A chủ yếu đo trầm cảm, chưa có RCT cho lo âu chuyên biệt tại VN.')
Pred('5. Can thiệp ĐA CẤP: cá nhân + GIA ĐÌNH + cộng đồng — kèm BOOSTER SESSIONS dài hạn để '
     'tránh fade-out (bài học Happy House).')

H('4.5. Kết luận (cập nhật v3)', level=2)
Pred('Việt Nam đã có RCT cluster đầu tiên về can thiệp SKTT học đường cho VTN (Happy House, '
     'Tran et al. 2023, Cambridge Prisms: Global Mental Health) — lấp khoảng trống "0 RCT VN" '
     'mà v2 đã nhận định sai. Tuy nhiên, khoảng trống vẫn còn lớn: Happy House là UNIVERSAL '
     '(toàn lớp), hiệu ứng nhỏ (d = 0,11) và FADE-OUT ở 6 tháng; đo TRẦM CẢM chính, không có '
     'lo âu chuyên biệt; chỉ lớp 10 Hà Nội. Bằng chứng quốc tế (Châu Á) cho thấy CBT (cá nhân, '
     'nhóm, internet) là phương pháp hiệu quả nhất; mô hình Sri Lanka (GV cung cấp) khả thi '
     'cho LMIC; iCBT phù hợp VTN số hoá; PHẢI tránh universal interventions chỉ với hiệu ứng '
     'nhỏ; cần tích hợp can thiệp GIA ĐÌNH (phát hiện Dong 2025) và BOOSTER SESSIONS dài hạn. '
     'Đề cương Việt Nam nên thiết kế can thiệp TARGETED 12 tuần đa thành phần + booster ở '
     '3 và 6 tháng — bổ sung cho Happy House đã có.',
     bold=True, italic=True)

# ============================================================
# PHẦN V — TÀI LIỆU THAM KHẢO MỚI (v3 bổ sung)
# ============================================================
doc.add_page_break()
Hred('TÀI LIỆU THAM KHẢO BỔ SUNG V3 (các bài mới so với v2)', level=1)

Pred('1. Tran, T.D., Nguyen, H., Shochet, I., Nguyen, N., La, N., Wurfl, A., Orr, J., '
     'Nguyen, H., Stocker, R., & Fisher, J. (2023). School-based universal mental health '
     'promotion intervention for adolescents in Vietnam: Two-arm, parallel, controlled trial. '
     'Cambridge Prisms: Global Mental Health, 10, e58. '
     'https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/')

Pred('2. Nguyen, N., Tran, T.D., Shochet, I., Nguyen, H., Fisher, J., et al. (2022). Adaptation '
     'of a school-based mental health program for adolescents in Vietnam. PLOS ONE, 17(8), '
     'e0271959. https://pubmed.ncbi.nlm.nih.gov/35925878/')

Pred('3. Studsgaard, E. et al. (2025). Developing an Internet-Based Cognitive Behavioral '
     'Therapy Intervention for Adolescents With Anxiety Disorders: Design, Usability, and '
     'Initial Evaluation of the CoolMinds Intervention. JMIR Formative Research, 9, e66966. '
     'https://pmc.ncbi.nlm.nih.gov/articles/PMC12015348/')

Pred('4. Baumgart, N. et al. (2025). Integrating Videoconferencing Therapist Guidance Into '
     'Stepped Care Internet-Delivered Cognitive Behavioral Therapy for Child and Adolescent '
     'Anxiety: Noninferiority Randomized Controlled Trial. JMIR Mental Health, 12, e57405. '
     'https://mental.jmir.org/2025/1/e57405')

Pred('5. [Author group TBD] (2025). School nurse–delivered anxiety interventions for '
     'adolescents: A systematic review and meta-analysis. Cogent Psychology (Taylor & Francis), '
     'published 18/06/2025. https://www.tandfonline.com/doi/full/10.1080/23311908.2025.2519529')

Pred('6. Imondi, G., Caron, E.B., Blanchard, L., Butler, E., & Ginsburg, G.S. (2025). Impact '
     'of the Child Anxiety Learning Modules (CALM), a School Nurse-Delivered Intervention for '
     'Anxiety, on Elementary Students School Outcomes. NASN School Nurse, 40. '
     'DOI 10.1177/10598405251328369')

Pred('7. [VR meta-analysis group] (2025). Effectiveness of virtual reality therapy in the '
     'treatment of anxiety disorders in adolescents and adults: a systematic review and '
     'meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, 1553290. '
     'https://pmc.ncbi.nlm.nih.gov/articles/PMC11904249/')

Pred('8. [JAACAP umbrella group] (2025). Systematic Umbrella Review and Meta-Meta-Analysis: '
     'Effectiveness of Physical Activity in Improving Depression and Anxiety in Children and '
     'Adolescents. Journal of the American Academy of Child & Adolescent Psychiatry. '
     'https://www.jaacap.org/article/S0890-8567(25)00208-4/fulltext')

Pred('9. Zapf, H., Stocker, J. (2024). A systematic review of the association between '
     'parent-child communication and adolescent mental health. JCPP Advances, 4. '
     'https://acamh.onlinelibrary.wiley.com/doi/10.1002/jcv2.12205')

Pred('10. Dao, T.N., Pham, T.S., Tran, T.N.T., Nguyen, T.N., Vo, T.H., Pho, T.V., Tran, H.T., '
     'Hoang, B.A., Tran, P.M., & Phan, T.M.N. (2025). Thực trạng tâm lý của sinh viên năm thứ '
     'tư Trường Đại học Y Hà Nội năm học 2023-2024. Tạp chí Nghiên cứu Y học, 187(02), tr. '
     '296–303. (= bài 59 của dự án)')

Pred('11. Thai, T.T., Vo Le, H.T., Nguyen, T.T., Dinh, N.V., Mai, X.L., Tran, H.T.T., Nguyen, '
     'N.B.T., Huynh, K.H.M., Nguyen, T.A.T., Bui, H.H.T., & Duong, M.C. (2025). Unmasking the '
     'burden of mental health symptoms and risk behaviors in Vietnamese adolescents: evidence '
     'from a multicenter cross-sectional study involving 2,631 high school students. Social '
     'Psychiatry and Psychiatric Epidemiology. DOI 10.1007/s00127-025-03043-7. (= bài 60)')

Pred('* Các bài từ v2 (23 bài gốc) giữ nguyên trong bản v2 — không liệt kê lại ở đây để tiết '
     'kiệm không gian.', italic=True, size=10)

# ============================================================
# PHẦN VI — TỪ KHOÁ ĐÃ DÙNG ĐỂ TÌM KIẾM (phụ lục)
# ============================================================
Hred('PHỤ LỤC A — TỪ KHOÁ ĐÃ DÙNG ĐỂ TÌM KIẾM WEB (v3)', level=1)
Pred('Sau đây là các từ khoá được rút trích từ các bài đã đọc và dùng để tìm kiếm bổ sung '
     'trên web (Google Scholar, PubMed, JMIR, tandfonline, Springer, Wiley):')
table(
    ['Nhóm từ khoá', 'Các từ khoá đã dùng'],
    [
        ['Thiết kế NC', 'randomized controlled trial, cluster RCT, meta-analysis, network meta-analysis, systematic review, umbrella review, feasibility trial, noninferiority trial'],
        ['Đối tượng', 'adolescent, children, young people, youth, high school, THCS, THPT, VTN, early adolescence, Vietnamese, Asian'],
        ['Rối loạn', 'anxiety disorder, social anxiety disorder (SAD), generalized anxiety disorder (GAD), subthreshold anxiety, panic, phobia, depression, stress, mental health'],
        ['Can thiệp cốt lõi', 'CBT, cognitive behavioral therapy, iCBT, internet-based CBT, DMHI, digital mental health, mobile CBT, smartphone app, VR, virtual reality exposure, CA-CBT, culturally adapted CBT, RAP-A, resilience training, mindfulness'],
        ['Bối cảnh', 'school-based, universal prevention, targeted prevention, indicated prevention, self-referral, stepped care, teacher-delivered, school nurse, MHST, PLACES, BESST, Happy House'],
        ['Tích hợp gia đình', 'parent-child communication, family intervention, parent training, parent-child interaction therapy, emotional support'],
        ['Địa lý', 'Vietnam, Southeast Asia, LMIC, Japan, Sri Lanka, China, UK, Australia, Denmark, United States'],
        ['Kết quả đo', 'Cohen d, SMD, SUCRA, Hedges g, OR, β coefficient, HAM-A, DASS-21, GAD-7, SIAS, SCAS'],
        ['Cơ chế', 'self-efficacy, resilience, optimism, perfectionism trap, exposure, habituation, attentional bias'],
    ],
    widths=[4.0, 11.5])

Pred('Các nguồn dữ liệu đã tra cứu: PubMed/PMC, JMIR journals (formative, mental health, '
     'pediatrics), tandfonline (Taylor & Francis), Springer/Nature, Wiley, Frontiers, Cambridge '
     'journals (Global Mental Health), JAMA Network, NEJM, BMC Psychiatry, PLOS ONE, PLOS Mental '
     'Health, ScienceDirect, Crossref, Semantic Scholar, Google Scholar.')

# ============================================================
# PHẦN VII — KẾT LUẬN VỀ BẢN V3
# ============================================================
Hred('KẾT LUẬN VỀ BẢN V3', level=2)
Pred('Báo cáo v3 đã cập nhật 11 điểm so với v2:', bold=True)
Pred('• Sửa 4 lỗi QA tìm được (VN25 tỷ lệ, QT44 tác giả, QT45 UMIN, QT46 tác giả).')
Pred('• Bổ sung 9 bài mới (1 bài VN: Happy House; 8 bài QT: CoolMinds, Stepped care VC, '
     'School nurse SR, CALM, VR MA, JAACAP PA umbrella, MIXCS Japan, Zapf parent-child, Brown & '
     'Carter UK).')
Pred('• Thêm 2 bài VN đã dịch mới (bài 59 TCNCYH, bài 60 Duong TPHCM).')
Pred('• Sửa khẳng định "0 RCT VN" thành "1 RCT VN (Happy House universal) + còn thiếu targeted".')
Pred('• Cập nhật bảng xếp hạng từ 13 → 15 phương pháp (thêm VR exposure + school nurse).')
Pred('• Cập nhật đề xuất 5 thành phần + booster sessions.')
Pred('• Thêm phụ lục từ khoá đã tìm kiếm trên web.')

P('---', italic=True, align='center')
P('Báo cáo v3 — 11/04/2026 — cập nhật từ v2 (10/04/2026).',
  italic=True, align='center', size=10)

doc.save(OUT)
print(f'Saved: {OUT}')
d2 = Document(OUT)
print(f'Total: {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables')
print(f'Word count estimate: {sum(len(p.text.split()) for p in d2.paragraphs)} words')
