# -*- coding: utf-8 -*-
"""Dump cac doan can sua trong LA v3_1."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_26052026.docx')
d = Document(LA)

# Paragraphs to dump
targets = [236, 247, 248, 249, 250, 261, 267, 268, 269, 271, 313, 315, 320, 331, 1368, 1373, 1375, 1454, 1481, 1499]
for t in targets:
    if t < len(d.paragraphs):
        print(f"\n--- para {t} ---")
        print(d.paragraphs[t].text)

# Find all "rất nặng" near Lam context, and "nam" context
print("\n=== TIM CAC CHO BAN DICH LAM CO THE BIA NAM>NU ===")
for i, p in enumerate(d.paragraphs):
    if 'Lâm' in p.text and ('nam' in p.text.lower() or 'nữ' in p.text.lower() or 'giới' in p.text.lower()):
        print(f"  para {i}: {p.text[:300]}")

# Check RLLATQ abbreviation usage
print("\n=== RLLATQ ===")
count_rllatq = sum(1 for p in d.paragraphs if 'RLLATQ' in p.text)
print(f"RLLATQ in paras: {count_rllatq}")
# Sample
for i, p in enumerate(d.paragraphs):
    if 'RLLATQ' in p.text:
        print(f"  para {i}: {p.text[:200]}")
        if i > 20: break

# Check tables for "tổng quát"
print("\n=== TABLES with 'tổng quát' ===")
for ti, tb in enumerate(d.tables):
    for ri, row in enumerate(tb.rows):
        for ci, cell in enumerate(row.cells):
            if 'tổng quát' in cell.text:
                print(f"  Table {ti+1} row {ri} cell {ci}: {cell.text[:150]}")
                break
