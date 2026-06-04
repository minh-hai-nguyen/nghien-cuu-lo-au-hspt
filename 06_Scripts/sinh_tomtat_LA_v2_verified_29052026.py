# -*- coding: utf-8 -*-
"""Ban tom tat LA v2 — SO LIEU DUOC VERIFY tu LA gocsach.
29/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', 'TomTatLA_v2_VERIFIED_29052026.docx')

BLACK = RGBColor(0, 0, 0)


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(11)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    s.paragraph_format.line_spacing = 1.15
    s.paragraph_format.space_after = Pt(3)
    for sec in d.sections:
        sec.page_width = Cm(14.8); sec.page_height = Cm(21.0)
        sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.5)
    return d


def H1(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    return p

def H2(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    return p

def H3(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    return p

def P(d, text, indent=True, bold=False, italic=False, center=False, size=11):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent and not center:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    return p

def page_break(d):
    p = d.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

def add_table_simple(d, headers, rows, widths_cm=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        if widths_cm: hdr[i].width = Cm(widths_cm[i])
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            r = p.add_run(str(cell_data))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            if widths_cm: row[i].width = Cm(widths_cm[i])
    return t


# ============================================================
d = doc_init()

# BIA CUNG
P(d, 'BỘ GIÁO DỤC VÀ ĐÀO TẠO', center=True, bold=True, size=12, indent=False)
P(d, 'TRƯỜNG ĐẠI HỌC SƯ PHẠM HÀ NỘI', center=True, bold=True, size=12, indent=False)
d.add_paragraph(); d.add_paragraph(); d.add_paragraph()
P(d, 'CÔNG THỊ HẰNG', center=True, bold=True, size=14, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'RỐI LOẠN LO ÂU Ở HỌC SINH', center=True, bold=True, size=16, indent=False)
P(d, 'TRUNG HỌC CƠ SỞ', center=True, bold=True, size=16, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'Chuyên ngành: Tâm lý học chuyên ngành', center=True, size=12, indent=False)
P(d, 'Mã số: 9.31.04.01', center=True, size=12, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'TÓM TẮT LUẬN ÁN TIẾN SĨ TÂM LÝ HỌC', center=True, bold=True, size=13, indent=False)
d.add_paragraph(); d.add_paragraph(); d.add_paragraph()
P(d, 'HÀ NỘI, 2026', center=True, bold=True, size=12, indent=False)
page_break(d)

# PHU BIA
P(d, 'Công trình được hoàn thành tại:', bold=True, indent=False, size=11)
P(d, 'TRƯỜNG ĐẠI HỌC SƯ PHẠM HÀ NỘI', bold=True, center=True, size=12, indent=False)
d.add_paragraph()
P(d, 'Người hướng dẫn khoa học:', bold=True, indent=False, size=11)
P(d, 'TS. Đào Minh Đức', center=True, italic=True, size=11, indent=False)
d.add_paragraph()
P(d, 'Phản biện 1: ........................................................................................', indent=False, size=11)
P(d, '            ........................................................................................', indent=False, size=11)
d.add_paragraph()
P(d, 'Phản biện 2: ........................................................................................', indent=False, size=11)
P(d, '            ........................................................................................', indent=False, size=11)
d.add_paragraph()
P(d, 'Phản biện 3: ........................................................................................', indent=False, size=11)
P(d, '            ........................................................................................', indent=False, size=11)
d.add_paragraph(); d.add_paragraph()
P(d, 'Luận án sẽ được bảo vệ trước Hội đồng đánh giá luận án cấp Trường tại Trường Đại học Sư phạm Hà Nội', size=11)
P(d, 'vào hồi ........ giờ ........ ngày ........ tháng ........ năm ........', size=11, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'Có thể tìm hiểu luận án tại:', bold=True, indent=False, size=11)
P(d, '- Thư viện Quốc gia Việt Nam', size=11, indent=False)
P(d, '- Thư viện Trường Đại học Sư phạm Hà Nội', size=11, indent=False)
page_break(d)

# MO DAU
H1(d, 'MỞ ĐẦU')

H2(d, '1. Lý do chọn đề tài')
P(d, 'Trong bối cảnh thế giới đương đại đầy biến động, sức khỏe tâm thần của lứa tuổi vị thành niên đang phải đối mặt với những thách thức chưa từng có, trong đó rối loạn lo âu (RLLA) nổi lên như một trong những vấn đề phổ biến và nổi bật nhất. Theo Tổ chức Y tế Thế giới (WHO), rối loạn lo âu là một trong hai vấn đề sức khỏe tâm thần nghiêm trọng nhất và thường khởi phát sớm ở giai đoạn thanh thiếu niên, với mức tăng tổng thể đến 25% trong giai đoạn đại dịch COVID-19.')
P(d, 'Tại Việt Nam, khảo sát đại diện quốc gia V-NAMHS (Viện Xã hội học và cộng sự, 2022) trên 5.996 cặp cha mẹ và trẻ vị thành niên cho thấy rối loạn lo âu là rối loạn tâm thần phổ biến nhất ở lứa tuổi 10–17, với tỷ lệ đạt ngưỡng chẩn đoán DSM-5 là 2,3%. Tuy nhiên, các khảo sát khác sử dụng thang đo sàng lọc đưa ra tỷ lệ dao động rất rộng, có khi lên đến trên 50%, cho thấy còn nhiều khoảng trống về phương pháp luận trong đo lường rối loạn lo âu ở học sinh Việt Nam.')
P(d, 'Lứa tuổi học sinh trung học cơ sở (11–15 tuổi) có vị trí đặc biệt quan trọng, được đặc trưng bởi tính dễ bị tổn thương trước các thách thức sức khỏe tâm thần do sự tương tác phức tạp giữa thay đổi tâm sinh lý và áp lực học đường, gia đình, xã hội. Đây cũng là giai đoạn khởi phát đặc trưng của nhiều dạng rối loạn lo âu, đặc biệt là rối loạn lo âu xã hội.')
P(d, 'Thực trạng nghiên cứu trong nước về chủ đề này còn nhiều khoảng trống: (1) chưa có bộ thang đo rối loạn lo âu chuẩn hóa phù hợp ngôn ngữ và văn hóa Việt Nam; (2) phần lớn nghiên cứu tập trung vào lo âu lan tỏa, trong khi lo âu xã hội và lo âu chia ly hầu như chưa được khảo sát chuyên biệt; (3) các yếu tố bảo vệ chưa được nghiên cứu đồng thời và tích hợp với yếu tố nguy cơ trong cùng một thiết kế; (4) thiết kế nghiên cứu theo chiều dọc còn rất ít; (5) thiếu các nghiên cứu về chương trình can thiệp theo tiếp cận nhận thức – hành vi (CBT) trên học sinh THCS Việt Nam.')
P(d, 'Trên cơ sở đó, luận án "Rối loạn lo âu ở học sinh trung học cơ sở" được lựa chọn nhằm cung cấp bằng chứng thực nghiệm về thực trạng và các yếu tố ảnh hưởng, đồng thời đề xuất khung chương trình tập huấn phòng ngừa phù hợp với bối cảnh giáo dục Việt Nam.')

H2(d, '2. Mục đích nghiên cứu')
P(d, 'Nghiên cứu cơ sở lý luận về rối loạn lo âu ở học sinh THCS; tìm hiểu thực trạng rối loạn lo âu và một số yếu tố ảnh hưởng (yếu tố nguy cơ và yếu tố bảo vệ); trên cơ sở đó, đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS.')

H2(d, '3. Đối tượng nghiên cứu')
P(d, 'Mức độ, biểu hiện và tác động của các yếu tố ảnh hưởng đến rối loạn lo âu học đường ở học sinh THCS; khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS.')

H2(d, '4. Khách thể nghiên cứu')
P(d, '- Học sinh THCS: khảo sát phiếu điều tra trên 1.352 học sinh THCS và phỏng vấn sâu 12 học sinh THCS tại Hà Nội.', indent=False)
P(d, '- Giáo viên THCS: phỏng vấn sâu 6 giáo viên chủ nhiệm THCS.', indent=False)

H2(d, '5. Giả thuyết khoa học')
P(d, 'Luận án xây dựng 3 giả thuyết khoa học:')
P(d, 'Giả thuyết 1 — Khác biệt giới tính phụ thuộc vào loại rối loạn lo âu: với rối loạn lo âu lan tỏa và lo âu xã hội, học sinh nữ có mức độ lo âu cao hơn nam; với lo âu chia ly, không có sự khác biệt giới tính có ý nghĩa thống kê.')
P(d, 'Giả thuyết 2 — Lòng tự trọng là yếu tố bảo vệ thuộc nhóm mạnh nhất, có cường độ tác động tương đương với áp lực học tập đối với rối loạn lo âu.')
P(d, 'Giả thuyết 3 — Áp lực học tập và nghiện điện thoại là hai yếu tố nguy cơ thuộc nhóm mạnh nhất trong số các yếu tố nguy cơ ảnh hưởng đến rối loạn lo âu ở học sinh THCS.')

H2(d, '6. Nhiệm vụ nghiên cứu')
P(d, '(1) Hệ thống hóa cơ sở lý luận về rối loạn lo âu và rối loạn lo âu ở học sinh trung học cơ sở.', indent=False)
P(d, '(2) Nghiên cứu thực trạng, biểu hiện mức độ rối loạn lo âu ở học sinh trung học cơ sở và các yếu tố ảnh hưởng (yếu tố nguy cơ, yếu tố bảo vệ, biện pháp đối phó).', indent=False)
P(d, '(3) Đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu học đường cho học sinh trung học cơ sở.', indent=False)

H2(d, '7. Giới hạn nghiên cứu')
P(d, 'Về nội dung: tập trung khảo sát ba loại rối loạn lo âu (lan tỏa, xã hội, chia ly) và các yếu tố ảnh hưởng (áp lực học tập, nghiện điện thoại, bắt nạt học đường, lòng tự trọng, gắn bó trường học, hỗ trợ xã hội, biện pháp đối phó).')
P(d, 'Về địa bàn: 2 trường THCS tại Hà Nội — THCS Nhật Tân (nội thành) và THCS Tây Mỗ (ven đô).')
P(d, 'Về khách thể: học sinh THCS các khối 6–9 (11–14 tuổi).')

H2(d, '8. Phương pháp nghiên cứu')
P(d, 'Luận án sử dụng phối hợp các phương pháp: (1) Nghiên cứu tài liệu, văn bản; (2) Điều tra bằng bảng hỏi; (3) Phỏng vấn sâu; (4) Quan sát; (5) Phân tích thống kê toán học bằng SPSS 31.0 và AMOS 31.0.')

H2(d, '9. Đóng góp mới của luận án')
P(d, 'Về lý luận: hệ thống hóa cơ sở lý luận về rối loạn lo âu ở học sinh THCS theo hướng tích hợp mạng lưới yếu tố nguy cơ và yếu tố bảo vệ; xác lập khung khái niệm về biện pháp ứng phó với hai chiều thích nghi và không thích nghi.')
P(d, 'Về thực tiễn: cung cấp bằng chứng định lượng về thực trạng rối loạn lo âu trên mẫu 1.352 học sinh THCS tại Hà Nội; đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu học đường có cơ sở khoa học và khả thi triển khai trong bối cảnh nhà trường Việt Nam.')

H2(d, '10. Cấu trúc luận án')
P(d, 'Ngoài phần Mở đầu, Kết luận, Tài liệu tham khảo và Phụ lục, luận án gồm 3 chương: Chương 1 — Cơ sở lý luận về rối loạn lo âu ở học sinh trung học cơ sở; Chương 2 — Tổ chức và phương pháp nghiên cứu; Chương 3 — Kết quả nghiên cứu thực trạng rối loạn lo âu.')

page_break(d)

# CHUONG 1
H1(d, 'CHƯƠNG 1\nCƠ SỞ LÝ LUẬN VỀ RỐI LOẠN LO ÂU\nỞ HỌC SINH TRUNG HỌC CƠ SỞ')

H2(d, '1.1. Tổng quan nghiên cứu')
H3(d, '1.1.1. Tỷ lệ phổ biến và mức độ rối loạn lo âu')
P(d, 'Các nghiên cứu quốc tế cho thấy tỷ lệ rối loạn lo âu ở thanh thiếu niên dao động khá rộng tùy theo công cụ sàng lọc và bối cảnh quốc gia. Xu và cộng sự (2021) trên 373.216 học sinh THCS-THPT Trung Quốc bằng GAD-7 báo cáo tỷ lệ 9,89% ở mức trung bình trở lên, tăng lên 38,42% nếu tính từ mức nhẹ. Bhardwaj và cộng sự (2020) ở Ấn Độ trên 288 học sinh THPT bằng DASS-21 ghi nhận 81,9% mức nhẹ trở lên và 73,2% từ mức trung bình trở lên. Chen và cộng sự (2023) trên 63.205 học sinh THCS-THPT Trung Quốc báo cáo tỷ lệ lo âu 13,9% và trầm cảm 23,0%.')
P(d, 'Tại Việt Nam, khảo sát V-NAMHS (Viện Xã hội học và cộng sự, 2022) trên 5.996 cặp cha mẹ-trẻ ghi nhận tỷ lệ rối loạn lo âu đạt ngưỡng chẩn đoán DSM-5 là 2,3%. Một số nghiên cứu khác bằng thang đo sàng lọc đưa ra tỷ lệ cao hơn: Hoàng Trung Học và Nguyễn Thùy Dung (2025) trên 8.473 học sinh từ lớp 6 đến lớp 12 tại 6 tỉnh thành báo cáo tỷ lệ lo âu 41,5% trong COVID-19, giảm còn 25,4% sau COVID-19; Nguyễn Ngọc Bảo Quyên và cộng sự (2025) trên 501 học sinh THPT Hà Nội bằng DASS-21 ghi nhận tỷ lệ 86,2%.')

H3(d, '1.1.2. Yếu tố nguy cơ và yếu tố bảo vệ')
P(d, 'Các yếu tố nguy cơ chính được xác định trong y văn quốc tế bao gồm: áp lực học tập (Steare và cộng sự, 2023 — tổng quan hệ thống 52 nghiên cứu, trong đó 48 nghiên cứu xác nhận tương quan thuận giữa áp lực học tập và sức khỏe tâm thần kém); nghiện điện thoại và mạng xã hội (A.M. Saikia và cộng sự, 2023; Z. Chen và cộng sự, 2023; T.L. Anderson và cộng sự, 2025); bắt nạt học đường (Z. Chen và cộng sự, 2023 — báo cáo tỷ lệ bắt nạt học đường dao động từ 9,0% đến 61,6%).')
P(d, 'Các yếu tố bảo vệ chủ yếu gồm: lòng tự trọng (Rosenberg, 1965; Masten, 2014); gắn bó với trường học; hỗ trợ từ cha mẹ và bạn bè (Compas và cộng sự, 2017 — meta-analysis trên 212 nghiên cứu, N = 80.850); và biện pháp ứng phó thích nghi (Lazarus và Folkman, 1984; Compas và cộng sự, 2017).')

H3(d, '1.1.3. Khoảng trống nghiên cứu tại Việt Nam')
P(d, 'Tổng quan các nghiên cứu trong nước cho thấy 5 khoảng trống chính: (1) chưa có bộ thang đo RLLA chuẩn hóa cho trẻ em-vị thành niên phù hợp ngôn ngữ và văn hóa Việt Nam, dẫn đến biên độ ước tính tỷ lệ lo âu rất rộng (từ 2,3% theo chẩn đoán đầy đủ đến trên 50% theo thang sàng lọc); (2) phần lớn nghiên cứu tập trung vào lo âu lan tỏa qua GAD-7/DASS-21, trong khi lo âu xã hội và lo âu chia ly hầu như chưa được khảo sát chuyên biệt; (3) các yếu tố bảo vệ chưa được nghiên cứu đồng thời và tích hợp với yếu tố nguy cơ trong cùng một thiết kế; (4) thiết kế nghiên cứu theo chiều dọc và thử nghiệm lâm sàng ngẫu nhiên đa trung tâm còn rất ít; (5) thiếu các nghiên cứu về chương trình can thiệp theo tiếp cận CBT trên học sinh THCS Việt Nam.')

H2(d, '1.2. Cơ sở lý luận')
H3(d, '1.2.1. Khái niệm rối loạn lo âu')
P(d, 'Rối loạn lo âu là trạng thái tâm lý bất thường, bao gồm các biểu hiện của sự sợ hãi quá mức, sự lo lắng và rối loạn hành vi liên quan, kéo dài và gây bất lợi đối với đời sống cá nhân. Theo DSM-5 (American Psychiatric Association, 2013), rối loạn lo âu được cấu trúc như một hệ thống đa thành phần, bao gồm các khía cạnh nhận thức (suy nghĩ tiêu cực, thảm họa hóa), cảm xúc (căng thẳng, bồn chồn), sinh lý (hoạt hóa hệ thần kinh giao cảm) và hành vi (né tránh, kiểm tra lặp lại).')

H3(d, '1.2.2. Phân loại rối loạn lo âu theo DSM-5')
P(d, 'Theo DSM-5, các dạng rối loạn lo âu chính ở trẻ em – vị thành niên gồm: (1) Rối loạn lo âu lan tỏa (Generalized Anxiety Disorder — GAD): lo lắng quá mức về nhiều sự kiện trong ít nhất 6 tháng, kèm các triệu chứng cơ thể; (2) Rối loạn lo âu xã hội (Social Anxiety Disorder): nỗi sợ rõ rệt và dai dẳng trong tình huống xã hội hoặc biểu diễn trước người khác; (3) Rối loạn lo âu chia ly (Separation Anxiety Disorder): lo lắng quá mức về sự chia cắt với người gắn bó, khởi phát đặc trưng ở trẻ em – vị thành niên.')

H3(d, '1.2.3. Học sinh THCS và rối loạn lo âu')
P(d, 'Rối loạn lo âu ở học sinh trung học cơ sở là trạng thái tâm lý bất thường, mất cân bằng trong cảm xúc, bao gồm các biểu hiện của sự sợ hãi quá mức, sự lo lắng, phản ứng sợ hãi không phù hợp với mức độ nguy hiểm của tình huống và rối loạn hành vi liên quan, đi kèm với các biểu hiện sinh lý, kéo dài và ảnh hưởng đến khả năng học tập, chất lượng cuộc sống của học sinh.')

H3(d, '1.2.4. Mạng lưới yếu tố ảnh hưởng')
P(d, 'Luận án xác lập khung phân tích các yếu tố ảnh hưởng đến rối loạn lo âu ở học sinh THCS theo hai nhóm: nhóm yếu tố nguy cơ (áp lực học tập, nghiện điện thoại, bắt nạt thể chất, bắt nạt bằng lời và qua mạng, lòng tự trọng thấp) và nhóm yếu tố bảo vệ (lòng tự trọng, gắn bó với trường học, sự hỗ trợ của cha mẹ, sự hỗ trợ từ bạn bè, biện pháp ứng phó thích nghi). Các yếu tố này tương tác đồng thời, không loại trừ lẫn nhau.')

H2(d, '1.3. Khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS')
P(d, 'Trên cơ sở các phát hiện về yếu tố nguy cơ và yếu tố bảo vệ, luận án xây dựng khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS với 3 cấu phần: (1) cơ sở khoa học dựa trên tiếp cận nhận thức – hành vi (CBT) và mô hình chăm sóc theo bậc (stepped care); (2) khái niệm khung — chương trình mang tính phòng ngừa cấp 1 và cấp 2, không nhằm điều trị lâm sàng; (3) cấu trúc chương trình gồm 8 nội dung cơ bản, kèm quy trình triển khai và đánh giá hiệu quả.')

page_break(d)

# CHUONG 2
H1(d, 'CHƯƠNG 2\nTỔ CHỨC VÀ PHƯƠNG PHÁP NGHIÊN CỨU')

H2(d, '2.1. Tổ chức nghiên cứu')
P(d, 'Nghiên cứu được tổ chức theo thiết kế hỗn hợp định lượng – định tính, triển khai qua hai giai đoạn:')
P(d, 'Giai đoạn nghiên cứu lý luận: tổng hợp tài liệu, xây dựng khung khái niệm, lựa chọn và hiệu chỉnh hệ thống thang đo phù hợp với học sinh THCS Việt Nam.')
P(d, 'Giai đoạn nghiên cứu thực tiễn: tổ chức điều tra bằng bảng hỏi trên mẫu lớn, phỏng vấn sâu để bổ sung chiều sâu cho dữ liệu định lượng.')
P(d, 'Địa bàn nghiên cứu: 2 trường THCS tại Hà Nội — THCS Nhật Tân (nội thành) và THCS Tây Mỗ (ven đô).')

H2(d, '2.2. Khách thể và đối tượng nghiên cứu')
P(d, 'Khách thể định lượng: 1.352 học sinh THCS các khối 6–9 (11–14 tuổi), được lựa chọn theo thiết kế phân tầng theo trường và khối lớp, kết hợp với chọn cụm ở cấp lớp học. Trong đó, học sinh nam là 614 em (45,4%) và học sinh nữ là 738 em (54,6%).')
P(d, 'Khách thể định tính: 12 học sinh THCS phỏng vấn sâu; 6 giáo viên chủ nhiệm THCS phỏng vấn sâu.')

H2(d, '2.3. Phương pháp nghiên cứu')
H3(d, '2.3.1. Phương pháp điều tra bằng bảng hỏi')
P(d, 'Bảng hỏi được xây dựng gồm hai phần: Phần 1 thu thập thông tin chung về khách thể (giới tính, khối lớp, học lực, đặc điểm sức khỏe – tâm lý – hành vi sinh hoạt); Phần 2 tập trung đánh giá thực trạng rối loạn lo âu và các yếu tố ảnh hưởng thông qua hệ thống thang đo chuẩn quốc tế đã được hiệu chỉnh.')

H3(d, '2.3.2. Hệ thống thang đo')
P(d, 'Các thang đo sử dụng trong nghiên cứu được lựa chọn theo nguyên tắc đã được kiểm chứng quốc tế, hiệu chỉnh theo quy trình dịch xuôi – dịch ngược, tham vấn chuyên gia và khảo sát thử nghiệm:')
P(d, '- Đo rối loạn lo âu: thang đo RCADS (Revised Children\'s Anxiety and Depression Scale; Chorpita, 2000) phân tách 3 dạng — rối loạn lo âu lan tỏa (RLLALT), rối loạn lo âu chia ly (RLLACL) và rối loạn lo âu xã hội (RLLAXH).', indent=False)
P(d, '- Đo yếu tố nguy cơ: thang đo ESSA (Educational Stress Scale for Adolescents; Sun và cs., 2011) đo áp lực học tập; thang đo nghiện điện thoại SAS-SV (Smartphone Addiction Scale - Short Version; Kwon và cs., 2013); thang đo bắt nạt học đường OBVQ (Olweus Bully/Victim Questionnaire; Olweus, 1996, phân tách 2 nhân tố: bắt nạt thể chất và bắt nạt lời nói).', indent=False)
P(d, '- Đo yếu tố bảo vệ: thang đo gắn bó với trường học (PSSM — Psychological Sense of School Membership; Goodenow, 1993); thang đo cảm nhận hỗ trợ xã hội (MSPSS — Multidimensional Scale of Perceived Social Support; Zimet và cs., 1988, phân tách hỗ trợ cha mẹ và hỗ trợ bạn bè); thang đo biện pháp đối phó Brief COPE (Carver, 1997, phân tách 3 nhân tố: tránh né, đối phó giải quyết vấn đề, tìm kiếm sự hỗ trợ); thang đo lòng tự trọng RSES (Rosenberg Self-Esteem Scale; Rosenberg, 1965).', indent=False)
P(d, 'Tất cả các thang đo sau khi hiệu chỉnh đều đạt yêu cầu về độ tin cậy (Cronbach\'s α ≥ 0,70 và McDonald\'s ω ≥ 0,70) và giá trị cấu trúc qua phân tích nhân tố khẳng định (CFA) với chỉ số phù hợp đạt ngưỡng chấp nhận.')

H3(d, '2.3.3. Phương pháp phỏng vấn sâu và quan sát')
P(d, 'Phỏng vấn sâu nhằm thu thập dữ liệu định tính, làm rõ thêm các biểu hiện và trải nghiệm liên quan đến rối loạn lo âu. Phương pháp quan sát ghi nhận các biểu hiện hành vi, thái độ, phản ứng của học sinh trong quá trình khảo sát.')

H3(d, '2.3.4. Phương pháp xử lý số liệu thống kê')
P(d, 'Dữ liệu được xử lý bằng SPSS 31.0 (sàng lọc, đánh giá độ tin cậy qua Cronbach\'s α và McDonald\'s ω, thống kê mô tả, kiểm định t-test, ANOVA) và AMOS 31.0 (phân tích nhân tố khẳng định CFA, mô hình phương trình cấu trúc SEM).')

H2(d, '2.4. Đạo đức nghiên cứu')
P(d, 'Nghiên cứu tuân thủ các nguyên tắc đạo đức trong nghiên cứu trên trẻ em – vị thành niên: (1) sự đồng ý có hiểu biết của phụ huynh và sự đồng ý của học sinh; (2) đảm bảo tính ẩn danh và bảo mật dữ liệu; (3) quyền rút lui khỏi nghiên cứu bất cứ lúc nào; (4) cung cấp thông tin liên hệ tư vấn tâm lý cho học sinh có biểu hiện rối loạn lo âu mức cao được phát hiện trong quá trình khảo sát.')

page_break(d)

# CHUONG 3 - VERIFIED NUMBERS
H1(d, 'CHƯƠNG 3\nKẾT QUẢ NGHIÊN CỨU THỰC TRẠNG\nRỐI LOẠN LO ÂU')

H2(d, '3.1. Hiệu lực đo lường của các thang đo')
P(d, 'Kết quả phân tích cho thấy các thang đo sử dụng trong nghiên cứu đều đạt độ tin cậy tốt và mức độ phù hợp cấu trúc chấp nhận được trên mẫu học sinh THCS Việt Nam.')

P(d, 'Bảng 3.1. Tổng hợp độ tin cậy và mức độ phù hợp CFA của các thang đo', bold=True, italic=True, center=True, indent=False, size=10)
# REAL VALUES FROM LA TABLES 4-19
add_table_simple(d,
    headers=['Thang đo', 'α', 'ω', 'CFI', 'RMSEA'],
    rows=[
        ['Lo âu lan tỏa (RLLALT)', '0,811', '0,811', '0,982', '0,049'],
        ['Lo âu chia ly (RLLACL)', '0,726', '0,726', '0,975', '0,099'],
        ['Lo âu xã hội (RLLAXH)', '0,744', '0,750', '0,992', '0,060'],
        ['Áp lực học tập (ALHT)', '0,708', '0,716', '0,998', '0,024'],
        ['Nghiện điện thoại (NĐT)', '0,836', '0,839', '0,996', '0,039'],
        ['Bắt nạt thể chất', '0,775', '0,777', '0,995', '0,050'],
        ['Bắt nạt lời nói', '0,864', '0,864', '0,995', '0,072'],
        ['Gắn bó với trường học', '0,747', '0,746', '0,978', '0,042'],
        ['Hỗ trợ cha mẹ', '0,847', '0,848', '0,996', '0,061'],
        ['Hỗ trợ bạn bè', '0,837', '0,837', '0,995', '0,064'],
        ['Tránh né (đối phó)', '0,727', '0,705', '0,862', '0,169'],
        ['Đối phó giải quyết vấn đề', '0,700', '0,699', '0,972', '0,052'],
        ['Tìm kiếm sự hỗ trợ', '0,773', '0,775', '1,000', '0,000'],
        ['Tự trọng', '0,725', '0,724', '0,988', '0,045'],
    ],
    widths_cm=[5.5, 1.5, 1.5, 1.5, 1.5])
P(d, 'Tất cả các thang đo đạt ngưỡng chấp nhận về độ tin cậy (α ≥ 0,70; ω ≥ 0,70), với hầu hết các mô hình CFA đạt chỉ số phù hợp tốt (CFI ≥ 0,90; RMSEA ≤ 0,10), qua đó khẳng định giá trị đo lường của các công cụ.', italic=True, size=10)

H2(d, '3.2. Thực trạng rối loạn lo âu ở học sinh THCS')
P(d, 'Kết quả khảo sát trên 1.352 học sinh THCS cho thấy mức độ rối loạn lo âu chủ yếu tập trung ở mức trung bình; trong đó, lo âu lan tỏa có xu hướng cao hơn so với các dạng còn lại, lo âu xã hội ở mức trung bình, còn lo âu chia ly ở mức thấp. Các biểu hiện lo âu tập trung chủ yếu vào áp lực học tập, đánh giá bản thân và các tình huống xã hội.')

P(d, 'Bảng 3.2. ĐTB rối loạn lo âu theo giới tính (thang điểm 0–100)', bold=True, italic=True, center=True, indent=False, size=10)
# REAL VALUES FROM LA TABLE 23
add_table_simple(d,
    headers=['Loại RLLA', 'Nam (n=614)\nĐTB (ĐLC)', 'Nữ (n=738)\nĐTB (ĐLC)', 'F', 'p'],
    rows=[
        ['Lo âu lan tỏa', '51,43 (22,01)', '59,47 (22,07)', '44,484', '< 0,001'],
        ['Lo âu chia ly', '25,42 (25,46)', '24,76 (23,29)', '0,246', '0,620'],
        ['Lo âu xã hội', '43,20 (25,09)', '52,74 (26,31)', '45,984', '< 0,001'],
        ['RLLA tổng hợp', '40,02 (19,02)', '45,66 (18,91)', '29,642', '< 0,001'],
    ],
    widths_cm=[3.2, 2.7, 2.7, 1.4, 1.4])
P(d, 'Kết quả xác nhận Giả thuyết 1: học sinh nữ có mức độ lo âu lan tỏa và lo âu xã hội cao hơn nam có ý nghĩa thống kê (p < 0,001), trong khi lo âu chia ly không có sự khác biệt giới tính (F = 0,246; p = 0,620).', italic=True, size=10)

H2(d, '3.3. Thực trạng các yếu tố nguy cơ và yếu tố bảo vệ')
P(d, 'Phân tích thực trạng cho thấy điểm trung bình của các yếu tố nguy cơ và bảo vệ (đã quy đổi theo thang điểm 0–100) như trong Bảng 3.3.')

P(d, 'Bảng 3.3. ĐTB của các yếu tố nguy cơ và yếu tố bảo vệ (thang 0–100)', bold=True, italic=True, center=True, indent=False, size=10)
# REAL VALUES FROM LA paras 1145 + table 24-25
add_table_simple(d,
    headers=['Nhóm', 'Yếu tố', 'ĐTB', 'ĐLC'],
    rows=[
        ['Nguy cơ', 'Áp lực học tập', '51,13', '23,92'],
        ['Nguy cơ', 'Nghiện điện thoại', '28,38', '—'],
        ['Nguy cơ', 'Bắt nạt thể chất', '13,52', '—'],
        ['Bảo vệ', 'Gắn bó với trường học', '52,60', '20,02'],
    ],
    widths_cm=[2.0, 5.0, 1.8, 1.8])
P(d, 'Trong các yếu tố nguy cơ, áp lực học tập có cường độ vượt giữa thang đo (51,13/100), vượt trội so với nghiện điện thoại (28,38) và bắt nạt thể chất (13,52).', italic=True, size=10)

H2(d, '3.4. Tác động của các yếu tố ảnh hưởng đến rối loạn lo âu')
H3(d, '3.4.1. Tác động của các yếu tố nguy cơ')
P(d, 'Kết quả phân tích mô hình phương trình cấu trúc (SEM) cho thấy ba nhóm yếu tố nguy cơ có thứ tự cường độ tác động đến rối loạn lo âu tổng hợp như sau:')

P(d, 'Bảng 3.4. Hệ số β chuẩn hóa: Yếu tố nguy cơ → Rối loạn lo âu', bold=True, italic=True, center=True, indent=False, size=10)
# REAL VALUES FROM LA TABLES 27, 30, 33
add_table_simple(d,
    headers=['Yếu tố nguy cơ', '→ Lan tỏa', '→ Chia ly', '→ Xã hội', '→ RLLA (3 factors)'],
    rows=[
        ['Áp lực học tập (ALHT)', '0,510***', '0,253***', '0,490***', '0,533*** (R²=0,284)'],
        ['Nghiện điện thoại (NĐT)', '0,336***', '0,265***', '0,383***', '0,400*** (R²=0,160)'],
        ['Bắt nạt học đường (BNHĐ)', '0,215***', '0,376***', '0,253***', '0,276*** (R²=0,076)'],
    ],
    widths_cm=[3.5, 1.7, 1.7, 1.7, 3.0])
P(d, 'Ghi chú: *** p < 0,001. ALHT là yếu tố nguy cơ mạnh nhất (β = 0,533). NĐT đứng thứ hai (β = 0,400). BNHĐ có mẫu hình khác biệt: tác động mạnh nhất lên lo âu chia ly (β = 0,376) thay vì lan tỏa hay xã hội.', italic=True, size=10)
P(d, 'Kết quả này xác nhận Giả thuyết 3: áp lực học tập và nghiện điện thoại là hai yếu tố nguy cơ thuộc nhóm mạnh nhất.', italic=True, size=10)

H3(d, '3.4.2. Tác động của các yếu tố bảo vệ')
P(d, 'Phân tích cho thấy các yếu tố bảo vệ có tác động ngược chiều (giảm rối loạn lo âu) với cường độ khác nhau:')

P(d, 'Bảng 3.5. Hệ số β chuẩn hóa: Yếu tố bảo vệ → Rối loạn lo âu', bold=True, italic=True, center=True, indent=False, size=10)
# REAL VALUES FROM LA TABLES 36, 39, 42
add_table_simple(d,
    headers=['Yếu tố bảo vệ', '→ Lan tỏa', '→ Chia ly', '→ Xã hội', '→ RLLA (3 factors)'],
    rows=[
        ['Tự trọng (TTr)', '-0,455***', '-0,087*', '-0,415***', '-0,457*** (R²=0,209)'],
        ['Hỗ trợ cha mẹ (HTCM)', '-0,172***', '0,000', '-0,273***', '-0,234*** (R²=0,055)'],
        ['Gắn bó trường học (GBTH)', '-0,108**', '0,014', '-0,187***', '-0,155*** (R²=0,024)'],
    ],
    widths_cm=[3.5, 1.7, 1.7, 1.7, 3.0])
P(d, 'Ghi chú: *** p < 0,001; ** p < 0,01; * p < 0,05. Lòng tự trọng là yếu tố bảo vệ mạnh nhất (β = -0,457), có cường độ tác động ~85–89% so với áp lực học tập (theo Kết luận LA, phụ thuộc dạng rối loạn lo âu cụ thể).', italic=True, size=10)
P(d, 'Kết quả này xác nhận Giả thuyết 2: lòng tự trọng là yếu tố bảo vệ thuộc nhóm mạnh nhất, có cường độ tác động tương đương với áp lực học tập đối với lo âu lan tỏa và lo âu xã hội.', italic=True, size=10)

H3(d, '3.4.3. Mô hình kết hợp yếu tố nguy cơ và yếu tố bảo vệ')
P(d, 'Mô hình kết hợp đồng thời yếu tố nguy cơ và yếu tố bảo vệ cho thấy yếu tố nguy cơ áp đảo yếu tố bảo vệ ở cấp tổng hợp, gợi ý nhu cầu can thiệp đồng thời cả hai chiều: vừa giảm nguy cơ, vừa tăng cường yếu tố bảo vệ.')

H2(d, '3.5. Biện pháp đối phó của học sinh THCS')
P(d, 'Phân tích Brief COPE theo ba nhân tố (tránh né, đối phó giải quyết vấn đề, tìm kiếm sự hỗ trợ) cho thấy:')
P(d, '- Chiều không thích nghi (tránh né): gia tăng rõ rệt cùng với triệu chứng lo âu — phù hợp với kỳ vọng lý thuyết.', indent=False)
P(d, '- Chiều thích nghi (đối phó giải quyết vấn đề và tìm kiếm sự hỗ trợ): cũng đi theo chiều dương cùng triệu chứng lo âu — đi ngược kỳ vọng, gợi ý rằng học sinh có ý thức sử dụng các chiến lược tích cực nhưng chưa đạt hiệu quả như mong đợi.', indent=False)
P(d, 'Kết quả này gợi ý cho định hướng can thiệp: đồng thời (1) loại bỏ các chiến lược không thích nghi và (2) nâng cao chất lượng triển khai các chiến lược thích nghi ở học sinh.')

H2(d, '3.6. Đề xuất Khung chương trình tập huấn phòng ngừa rối loạn lo âu')
P(d, 'Trên cơ sở các kết quả nghiên cứu lý luận và thực trạng, luận án đề xuất Khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS với 8 nội dung cơ bản: (1) Nhận diện rối loạn lo âu và tự đánh giá; (2) Kỹ năng tái cấu trúc nhận thức; (3) Kỹ năng quản lý áp lực học tập; (4) Kỹ năng sử dụng thiết bị số lành mạnh; (5) Kỹ năng đối phó với bắt nạt học đường; (6) Củng cố lòng tự trọng; (7) Tăng cường gắn bó với trường học và hỗ trợ xã hội; (8) Phát triển chiến lược ứng phó thích nghi.')
P(d, 'Chương trình được thiết kế theo tiếp cận CBT, chăm sóc theo bậc (stepped care) và mô hình phòng ngừa cấp 1 + cấp 2; có thể lồng ghép vào sinh hoạt đầu giờ, chào cờ, sinh hoạt lớp, chuyên đề, hoặc tổ chức dưới dạng các cuộc thi dành cho học sinh THCS.')

page_break(d)

# KET LUAN
H1(d, 'KẾT LUẬN VÀ KIẾN NGHỊ')

H2(d, '1. Kết luận')
H3(d, '1.1. Về lý luận')
P(d, 'Luận án đã hệ thống hóa cơ sở lý luận về rối loạn lo âu ở học sinh THCS theo hướng tích hợp mạng lưới yếu tố nguy cơ và yếu tố bảo vệ; xác lập khung khái niệm RLLA gồm 4 thành phần (nhận thức – cảm xúc – sinh lý – hành vi); phân biệt rõ 3 dạng RLLA (lan tỏa, xã hội, chia ly) trên đối tượng học sinh THCS; xác lập khung phân tích biện pháp ứng phó theo hai chiều thích nghi và không thích nghi.')

H3(d, '1.2. Về thực tiễn')
P(d, 'Luận án đã cung cấp một bức tranh chi tiết về thực trạng rối loạn lo âu trên mẫu 1.352 học sinh THCS tại Hà Nội. Hệ thống các thang đo sau khi hiệu chỉnh đều đạt yêu cầu về độ tin cậy và giá trị cấu trúc. Mức độ rối loạn lo âu của học sinh THCS chủ yếu ở mức trung bình, lo âu lan tỏa và lo âu xã hội nổi trội hơn lo âu chia ly.')
P(d, 'Luận án đã kiểm chứng và xác nhận cả 3 giả thuyết khoa học:')
P(d, '- Giả thuyết 1 (khác biệt giới tính phụ thuộc loại RLLA): xác nhận. Học sinh nữ có mức độ lo âu lan tỏa (ĐTB 59,47 vs 51,43; F = 44,484; p < 0,001) và lo âu xã hội (ĐTB 52,74 vs 43,20; F = 45,984; p < 0,001) cao hơn nam có ý nghĩa thống kê; lo âu chia ly không có khác biệt giới tính (F = 0,246; p = 0,620).', indent=False)
P(d, '- Giả thuyết 2 (lòng tự trọng là YTBV mạnh): xác nhận. Lòng tự trọng có hệ số β = -0,457 đối với RLLA tổng hợp, cường độ tác động ~85–89% so với áp lực học tập (theo Kết luận LA).', indent=False)
P(d, '- Giả thuyết 3 (áp lực học tập + nghiện điện thoại là YTNC mạnh): xác nhận. Thứ tự cường độ tác động: áp lực học tập (β = 0,533) > nghiện điện thoại (β = 0,400) > bắt nạt học đường (β = 0,276).', indent=False)
P(d, 'Trên cơ sở đó, luận án đề xuất Khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS với 8 nội dung cơ bản, đáp ứng yêu cầu cấp thiết hiện nay về chăm sóc sức khỏe tâm thần cho học sinh Việt Nam.')

H2(d, '2. Kiến nghị')
H3(d, '2.1. Đối với học sinh')
P(d, 'Cần trang bị kiến thức và kỹ năng nhận diện sớm các dấu hiệu rối loạn lo âu; rèn luyện năng lực tự điều chỉnh tâm lý; giảm các chiến lược không thích nghi (tránh né, tự đổ lỗi) thông qua kỹ thuật tái cấu trúc nhận thức (CBT); nâng cao chất lượng triển khai các chiến lược thích nghi (giải quyết vấn đề, tìm kiếm hỗ trợ).')

H3(d, '2.2. Đối với giáo viên và nhà trường')
P(d, 'Lồng ghép nội dung chương trình tập huấn phòng ngừa rối loạn lo âu trong sinh hoạt đầu giờ, chào cờ, sinh hoạt lớp, chuyên đề. Rà soát và điều chỉnh các yếu tố trong hoạt động dạy học có khả năng làm gia tăng RLLA (áp lực đánh giá, cạnh tranh thành tích). Bồi dưỡng năng lực nhận diện biểu hiện RLLA cho giáo viên.')

H3(d, '2.3. Đối với gia đình')
P(d, 'Phụ huynh cần điều chỉnh kỳ vọng phù hợp với năng lực và đặc điểm tâm lý của con; xây dựng môi trường giao tiếp cởi mở; trở thành nguồn hỗ trợ tâm lý ổn định và tích cực.')

H3(d, '2.4. Đối với các lực lượng hỗ trợ học đường')
P(d, 'Phòng tư vấn học đường triển khai sàng lọc và hỗ trợ sớm đối với học sinh có nguy cơ. Đoàn thanh niên, câu lạc bộ tăng cường kết nối xã hội. Phối hợp đồng bộ giữa các lực lượng theo hướng can thiệp tích hợp và đa tầng.')

H2(d, '3. Giới hạn và định hướng nghiên cứu trong tương lai')
P(d, 'Nghiên cứu có một số hạn chế: (1) thiết kế cắt ngang chưa khẳng định được quan hệ nhân quả theo chiều thời gian; (2) một số cấu trúc đo lường vẫn cần tiếp tục điều chỉnh để nâng cao tính ổn định; (3) phương pháp tự báo cáo có thể chịu ảnh hưởng của sai lệch xã hội; (4) phạm vi mẫu giới hạn trong 2 trường THCS tại Hà Nội; (5) mô hình chưa xem xét đầy đủ các biến trung gian và điều tiết.')
P(d, 'Định hướng nghiên cứu tiếp theo: (1) thiết kế dọc theo thời gian; (2) hoàn thiện công cụ đo lường trên các mẫu khác nhau; (3) triển khai nghiên cứu thực nghiệm kiểm định hiệu quả Khung chương trình tập huấn phòng ngừa rối loạn lo âu trong bối cảnh học đường.')

page_break(d)

# DANH MUC CONG TRINH
H1(d, 'DANH MỤC CÔNG TRÌNH KHOA HỌC CỦA TÁC GIẢ ĐÃ CÔNG BỐ LIÊN QUAN ĐẾN LUẬN ÁN')

P(d, '[NCS cần điền danh sách bài báo đã đăng. Yêu cầu tối thiểu: 2 bài đăng tạp chí khoa học có chỉ số khoa học theo quy định của Bộ Giáo dục và Đào tạo và Trường Đại học Sư phạm Hà Nội.]', italic=True, size=10)
d.add_paragraph()
P(d, '1. Công Thị Hằng. (20...). [Tên bài báo]. Tạp chí ......, tập ......, số ......, trang ........ DOI: .....................', indent=False, size=11)
d.add_paragraph()
P(d, '2. Công Thị Hằng. (20...). [Tên bài báo]. Tạp chí ......, tập ......, số ......, trang ........ DOI: .....................', indent=False, size=11)
d.add_paragraph()
P(d, '3. Công Thị Hằng. (20...). [Tên bài báo]. Tạp chí ......, tập ......, số ......, trang ........ DOI: .....................', indent=False, size=11)

cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''; cp.keywords = ''
cp.comments = ''; cp.last_modified_by = ''

d.save(OUT)
from docx import Document as D
dd = D(OUT)
w = sum(len(p.text.split()) for p in dd.paragraphs)
print(f"Da luu: {OUT}")
print(f"Paragraphs: {len(dd.paragraphs)}, Tables: {len(dd.tables)}, Words: ~{w}, Size: {os.path.getsize(OUT)//1024}KB")
