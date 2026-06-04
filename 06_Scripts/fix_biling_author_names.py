# -*- coding: utf-8 -*-
"""Fix Vietnamese author name ordering in bilingual refs.

Rule: Vietnamese author names use surname first: Đặng Hoàng Minh (not Hoàng Minh Đặng).
For authors where Vietnamese ordering is unclear, keep Western initial format (L.T. Trung).
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

OUT = '03_Ban-dich/VN002_VNAMHS_2022_National_FULL.docx'
d = Document(OUT)

# Fix author ordering
FIXES = [
    # Hoàng Minh Đặng (wrong English-order) → Đặng Hoàng Minh (correct VN order)
    ('Hoàng Minh Đặng', 'Đặng Hoàng Minh'),
    # "Mai Đỗ" was guess — revert to "Đỗ Mai D." to keep safer
    ('Mai Đỗ, Phạm N.N.K.', 'Đỗ Mai, Phạm N.N.K.'),
    # Khỏi fabrication: "Đỗ L. Sang" → "D.L. Sang" (safer - don't guess VN name)
    ('Đỗ L. Sang', 'D.L. Sang'),
    # "Lê T. Trung" → "L.T. Trung" (keep initials safe)
    ('Lê T. Trung', 'L.T. Trung'),
    # Thái T. Trúc → keep T.T. Trúc to avoid fabrication
    ('Thái T. Trúc', 'T.T. Trúc'),
]

changes = 0
for p in d.paragraphs:
    t = p.text
    new = t
    for f, r in FIXES:
        if f in new:
            new = new.replace(f, r)
    if new != t:
        for run in p.runs: run.text = ''
        if p.runs: p.runs[0].text = new
        changes += 1

d.save(OUT)
print(f'Author ordering fixes applied: {changes}')

# Verify
d2 = Document(OUT)
for p in d2.paragraphs:
    if 'Đặng Hoàng Minh' in p.text:
        print(f'  ✓ "Đặng Hoàng Minh" correctly used: {p.text[:80]}')
    if 'Hoàng Minh Đặng' in p.text:
        print(f'  ⚠ "Hoàng Minh Đặng" still present (should not): {p.text[:80]}')
