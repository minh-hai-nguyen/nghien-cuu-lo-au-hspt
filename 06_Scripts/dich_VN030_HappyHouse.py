# -*- coding: utf-8 -*-
"""Tao ban dich day du VN030 Happy House + cap nhat tom tat."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DICH_DIR = os.path.join(ROOT, '03_Ban-dich')
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name='Times New Roman'; s.font.size=Pt(12)
    s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.5
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3); sec.right_margin=Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
def Pred(doc, text, bold=False):
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(12); r.bold=bold
    r.font.color.rgb=RGBColor(0xCC,0,0)
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)

# ============================================================
# BẢN DỊCH ĐẦY ĐỦ VN030 — Happy House Tran et al. 2023
# ============================================================
doc = make_doc()
P(doc, 'Link bài báo gốc: https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/', size=10)
P(doc, 'Open Access — Cambridge University Press', size=10, italic=True)

H(doc, 'HAPPY HOUSE — Can thiệp thúc đẩy sức khoẻ tâm thần phổ quát tại trường cho vị thành niên Việt Nam: Thử nghiệm đối chứng hai nhánh song song', level=1)
p = doc.add_paragraph()
r = p.add_run('"School-based universal mental health promotion intervention for adolescents in '
              'Vietnam: Two-arm, parallel, controlled trial"')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

H(doc, 'THÔNG TIN THƯ MỤC', 2)
table(doc,
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Thach Duc Tran, Huong Nguyen, Ian Shochet, Nga Nguyen, Nga La, Astrid Wurfl, Jayne Orr, Hau Nguyen, Ruby Stocker, Jane Fisher'],
        ['Đơn vị', 'Monash University (Melbourne, Úc); Hanoi University of Public Health (Hà Nội, VN); Queensland University of Technology (Brisbane, Úc)'],
        ['Tạp chí', 'Cambridge Prisms: Global Mental Health, vol. 10, e69 (10/2023)'],
        ['Link PMC', 'https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/'],
        ['Đăng ký thử nghiệm', 'ACTRN12620000088943'],
        ['Loại NC', 'Thử nghiệm đối chứng hai nhánh song song (cluster non-randomized)'],
        ['Mẫu', '1.084 HS lớp 10 (96,1 % tỷ lệ tham gia từ 1.128 HS đủ điều kiện); 531 can thiệp + 552 đối chứng; 8 trường THPT Hà Nội (4+4)'],
        ['Tuổi', '15–16 tuổi'],
        ['Đặc điểm mẫu', 'Nữ 58,9 % (đối chứng) vs 62,3 % (can thiệp); đô thị ~50 %; sống cùng cả cha mẹ ~89 %; CESD-R ban đầu 11,4 ± 12,2 vs 12,0 ± 12,0 (p = 0,729)'],
        ['Tài trợ', 'NHMRC Úc (GNT1158429) + NAFOSTED Việt Nam (NHMRC.108.01-2018.02)'],
        ['Đạo đức', 'Monash HREC + HUPH IRB + QUT OREC'],
    ],
    widths=[4.0, 11.5])

H(doc, 'TÓM TẮT', 2)
P(doc, 'Bối cảnh: Các can thiệp tâm lý dựa trên trường học đã cho thấy hiệu quả trong giảm '
       'các triệu chứng trầm cảm ở vị thành niên tại các nước thu nhập cao, nhưng bằng chứng '
       'từ các nước thu nhập thấp – trung bình (LMIC) còn hạn chế. Chương trình Tài nguyên Vị '
       'thành niên (Resourceful Adolescent Program — RAP) là một trong những can thiệp trường '
       'học dựa trên bằng chứng được nghiên cứu nhiều nhất. Nghiên cứu này đánh giá hiệu quả '
       'của chương trình "Happy House" — phiên bản thích ứng văn hoá Việt Nam của RAP — trong '
       'thúc đẩy sức khoẻ tâm thần của vị thành niên Việt Nam.')

P(doc, 'Phương pháp: Thử nghiệm đối chứng hai nhánh song song, không ngẫu nhiên cấp HS. 8 '
       'trường THPT tại Hà Nội được phân (4 can thiệp + 4 đối chứng) dựa trên sự đồng ý của '
       'ban giám hiệu. 1.084 học sinh lớp 10 được tuyển (96,1 % tỷ lệ) với 531 nhóm can thiệp '
       'và 552 nhóm đối chứng. Nhóm can thiệp nhận chương trình Happy House gồm 6 buổi, mỗi '
       'buổi 90 phút, trong 6 tuần, thay thế môn Giáo dục Công dân. Nhóm đối chứng tiếp tục '
       'chương trình học chính quy. Đánh giá tại ba thời điểm: trước can thiệp, 2 tuần sau, '
       '6 tháng sau. Kết cục chính: triệu chứng trầm cảm theo Center for Epidemiologic Studies '
       'Depression Scale (CESD-R, 20 mục, thang 0–80, ngưỡng ≥ 16 cho triệu chứng lâm sàng có '
       'ý nghĩa). Kết cục phụ: phúc lợi tâm thần (MHC-SF), tự hiệu quả ứng phó (CSES), kết nối '
       'trường học. Phân tích hỗn hợp mixed-effects models kiểm soát cho hiệu ứng cụm và đặc '
       'điểm ban đầu. Phân tích theo chủ đích điều trị (ITT).')

P(doc, 'Kết quả: Thời điểm 2 tuần sau can thiệp, tỷ lệ học sinh có CESD-R ≥ 16 là 28,6 % '
       '(155/542) ở nhóm đối chứng so với 23,9 % (125/523) ở nhóm can thiệp. Tỷ suất chênh đã '
       'điều chỉnh aOR = 0,56 (KTC 95 %: 0,36 – 0,88; p = 0,011), với Cohen d = 0,11. Nhóm '
       'can thiệp cũng có cải thiện về phúc lợi tâm lý (effect size 0,13; p < 0,05) và tự '
       'hiệu quả ứng phó (effect sizes 0,17 – 0,26) ở mọi ba khía cạnh (tập trung vào vấn đề, '
       'tập trung vào cảm xúc, tìm kiếm hỗ trợ xã hội). Tại 6 tháng, tỷ lệ CESD-R ≥ 16 là '
       '29,2 % vs 26,3 % với aOR = 0,75 (KTC 0,51 – 1,09; p = 0,132) — không còn ý nghĩa '
       'thống kê. Tuy nhiên, tự hiệu quả ứng phó tập trung vào vấn đề và tìm kiếm hỗ trợ xã '
       'hội vẫn DUY TRÌ ý nghĩa ở 6 tháng.')

P(doc, 'Kết luận: Chương trình Happy House cho thấy tiềm năng lớn để tích hợp vào chương '
       'trình học hiện có nhằm giảm các vấn đề sức khoẻ tâm thần ở vị thành niên Việt Nam. '
       'Tuy nhiên, hiệu quả trên trầm cảm có FADE-OUT ở 6 tháng, một phần do đại dịch COVID-19 '
       'làm gián đoạn thời gian theo dõi. Các tác giả đề xuất bổ sung BOOSTER SESSIONS để duy '
       'trì hiệu quả, phân tích chi phí – hiệu quả, và giải quyết các rào cản về bản quyền '
       'chương trình RAP.')

H(doc, 'GIỚI THIỆU', 2)
P(doc, 'Sức khoẻ tâm thần là một vấn đề y tế công cộng ngày càng gia tăng ở trẻ vị thành '
       'niên toàn cầu. Tại Việt Nam, các nghiên cứu trước đã báo cáo tỷ lệ triệu chứng trầm '
       'cảm lâm sàng khoảng 23 % ở HS THPT. Các can thiệp trường học phổ quát — nhắm vào toàn '
       'bộ HS trong lớp, không chỉ HS có triệu chứng — được ưa chuộng tại LMIC vì khả năng '
       'tiếp cận rộng, giảm kỳ thị, và có thể tích hợp vào hệ thống giáo dục hiện có.')

P(doc, 'Chương trình RAP (Resourceful Adolescent Program) được phát triển tại Úc, kết hợp các '
       'thành phần của CBT và liệu pháp tương tác cá nhân (IPT). Mục tiêu: giúp HS phát triển '
       'kỹ năng nhận thức, cảm xúc, và xã hội để ứng phó với các thách thức phát triển. Các '
       'thử nghiệm RAP tại Úc và New Zealand cho thấy hiệu quả bền vững 10–18 tháng sau can '
       'thiệp. Tuy nhiên, việc áp dụng RAP tại LMIC còn hạn chế do khác biệt văn hoá, ngôn '
       'ngữ, và cấu trúc lớp học (lớp học châu Á thường đông hơn, 40–45 HS so với 16 HS ở Úc).')

P(doc, 'Nghiên cứu này nhằm: (1) thích ứng RAP cho bối cảnh văn hoá + giáo dục Việt Nam — '
       'tạo ra chương trình "Happy House"; (2) đánh giá hiệu quả của Happy House trong giảm '
       'triệu chứng trầm cảm và cải thiện phúc lợi tâm thần của HS lớp 10 tại Hà Nội.')

H(doc, 'PHƯƠNG PHÁP', 2)
H(doc, 'Thiết kế nghiên cứu', 3)
P(doc, 'Thử nghiệm đối chứng hai nhánh song song, không ngẫu nhiên. Phân bổ cấp trường (cluster '
       'non-randomized). Do các yêu cầu hành chính của Sở Giáo dục Hà Nội, 8 trường THPT được '
       'chọn dựa trên (a) sẵn sàng tham gia của ban giám hiệu, (b) cân bằng về đặc điểm kinh '
       'tế – xã hội, (c) địa lý (đô thị vs ngoại ô). 4 trường được phân vào nhóm can thiệp và '
       '4 trường vào nhóm đối chứng.')

H(doc, 'Mẫu nghiên cứu', 3)
P(doc, 'Từ 1.128 HS lớp 10 đủ điều kiện tại 8 trường, 1.084 HS (96,1 %) đồng ý tham gia. '
       'Phân bổ: 531 nhóm can thiệp (8–14 lớp/trường) + 552 nhóm đối chứng (8–14 lớp/trường). '
       'Tuổi 15–16. Đặc điểm ban đầu:')
table(doc,
    ['Đặc điểm', 'Đối chứng (n=552)', 'Can thiệp (n=531)', 'p'],
    [
        ['Giới nữ', '58,9 %', '62,3 %', '0,261'],
        ['Đô thị', '51,5 %', '48,2 %', '0,298'],
        ['Sống cùng cả cha mẹ', '89,9 %', '88,7 %', '0,524'],
        ['CESD-R ban đầu (TB ± SD)', '11,4 ± 12,2', '12,0 ± 12,0', '0,729'],
    ],
    widths=[5.0, 3.5, 3.5, 3.5])

H(doc, 'Can thiệp Happy House', 3)
P(doc, 'Happy House được thích ứng từ chương trình RAP Úc qua 4 giai đoạn: (1) dịch sang tiếng '
       'Việt + hiệu đính ngược; (2) thích ứng văn hoá với sự tham gia của các chuyên gia tâm '
       'lý học đường Việt Nam; (3) thử nghiệm pilot với 30 HS Hà Nội; (4) sửa đổi dựa trên '
       'phản hồi HS và giáo viên.')

P(doc, 'Cấu trúc cuối cùng: 11 buổi 45 phút của RAP gốc được tái cấu trúc thành 6 buổi 90 phút, '
       'thực hiện trong 6 tuần liền. Mỗi buổi Happy House thay thế 2 tiết Giáo dục Công dân '
       'liên tiếp. Nội dung chính của 6 buổi bao gồm:')
table(doc,
    ['Buổi', 'Chủ đề', 'Kỹ năng chính'],
    [
        ['1', 'Sức mạnh bản thân', 'Nhận diện điểm mạnh, xây dựng tự trọng'],
        ['2', 'Suy nghĩ tích cực', 'Tái cấu trúc nhận thức (CBT), dialogues nội tâm'],
        ['3', 'Quản lý cảm xúc', 'Kỹ thuật thư giãn, điều tiết cảm xúc'],
        ['4', 'Giải quyết vấn đề', 'Các bước giải quyết vấn đề có cấu trúc'],
        ['5', 'Mạng lưới hỗ trợ', 'Xây dựng quan hệ, tìm kiếm giúp đỡ (IPT)'],
        ['6', 'Ôn tập và cam kết', 'Tổng kết kỹ năng, cam kết tự thực hành'],
    ],
    widths=[1.0, 4.5, 10.0])

P(doc, 'Người cung cấp: Giáo viên Giáo dục Công dân của mỗi trường (có nền tảng tâm lý cơ bản) '
       'được đào tạo 3 ngày bởi các chuyên gia Monash + Hanoi University of Public Health. '
       'Mỗi buổi can thiệp có giáo viên chính + 1–2 trợ giảng để hỗ trợ lớp đông (40–45 HS).')

P(doc, 'Tỷ lệ tham gia: 95 % HS hoàn thành toàn bộ 6 buổi, 5 % vắng mặt 1 buổi. Fidelity '
       '100 % hoạt động được cung cấp đúng kế hoạch (tự đánh giá của giáo viên — COVID-19 '
       'ngăn cản quan sát trực tiếp từ nhóm Úc).')

H(doc, 'Công cụ đo', 3)
P(doc, 'Kết cục chính — CESD-R (Center for Epidemiologic Studies Depression Scale Revised): '
       '20 mục, thang 0–80, ngưỡng cắt ≥ 16 cho triệu chứng trầm cảm có ý nghĩa lâm sàng. '
       'Phiên bản tiếng Việt đã được validate trước đó.')

P(doc, 'Kết cục phụ:')
P(doc, '• Mental Health Continuum Short Form (MHC-SF): đo phúc lợi tâm thần theo 3 khía cạnh '
       '— cảm xúc (emotional), xã hội (social), tâm lý (psychological)')
P(doc, '• Coping Self-Efficacy Scale (CSES): 3 thang — giải quyết vấn đề (problem-focused), '
       'điều tiết cảm xúc (emotion-focused), tìm kiếm hỗ trợ xã hội (social support-seeking)')
P(doc, '• School Connectedness (5 mục, thang 0–20)')

P(doc, 'Thời điểm đo: (1) Ban đầu — 05–24/10/2020; (2) Sau can thiệp 2 tuần — 01–19/12/2020; '
       '(3) Theo dõi 6 tháng — 10–27/05/2021 (chuyển sang khảo sát trực tuyến do COVID-19).')

H(doc, 'Phân tích thống kê', 3)
P(doc, 'Phân tích theo chủ đích điều trị (intention-to-treat — ITT). Mixed-effects models '
       'với hiệu ứng cố định (trường, nhóm điều trị, thời điểm, tương tác nhóm × thời điểm) và '
       'hiệu ứng ngẫu nhiên (cụm lớp học). Kiểm soát cho các biến số ban đầu: tuổi, giới, đô '
       'thị, trình độ cha mẹ, điểm CESD-R ban đầu. Cohen d được tính cho kích thước hiệu ứng '
       'giữa nhóm. Hiệu chỉnh Bonferroni cho nhiều so sánh kết cục phụ.')

H(doc, 'KẾT QUẢ', 2)

H(doc, 'Tuyển mộ và mất mẫu', 3)
P(doc, '1.084 HS được tuyển (96,1 % tỷ lệ). Mất mẫu: 29 HS thiếu dữ liệu ban đầu (2,7 %), '
       '21 HS thiếu dữ liệu sau can thiệp (1,9 %), 20 HS thiếu dữ liệu 6 tháng (1,9 %). Tỷ '
       'lệ mất mẫu rất thấp so với các thử nghiệm trường học khác.')

H(doc, 'Kết cục chính — Trầm cảm (CESD-R)', 3)
table(doc,
    ['Thời điểm', 'Đối chứng (≥16)', 'Can thiệp (≥16)', 'aOR (KTC 95 %)', 'p', 'Cohen d'],
    [
        ['2 tuần sau', '155/542 (28,6 %)', '125/523 (23,9 %)', '0,56 (0,36–0,88)', '0,011', '0,11'],
        ['6 tháng sau', '157/537 (29,2 %)', '138/524 (26,3 %)', '0,75 (0,51–1,09)', '0,132', '—'],
    ],
    widths=[2.5, 3.2, 3.2, 3.2, 1.8, 1.6])

P(doc, 'Tại thời điểm 2 tuần sau can thiệp, nhóm can thiệp có tỷ lệ thấp hơn nhóm đối chứng '
       '4,7 điểm phần trăm về triệu chứng trầm cảm có ý nghĩa lâm sàng. aOR = 0,56 có nghĩa '
       'là khả năng trầm cảm lâm sàng của HS can thiệp chỉ bằng 56 % so với HS đối chứng '
       '(p = 0,011). Tuy nhiên tại 6 tháng, tác dụng không còn ý nghĩa thống kê (aOR = 0,75, '
       'p = 0,132) — hiện tượng FADE-OUT.', italic=True)

H(doc, 'Kết cục phụ', 3)
table(doc,
    ['Kết cục phụ', '2 tuần sau (d)', '6 tháng sau (d)'],
    [
        ['MHC-SF Phúc lợi cảm xúc', 'Không ý nghĩa', 'Không ý nghĩa'],
        ['MHC-SF Phúc lợi xã hội', 'Không ý nghĩa', 'Không ý nghĩa'],
        ['MHC-SF Phúc lợi tâm lý', '0,13 *', 'Không ý nghĩa'],
        ['CSES Ứng phó cảm xúc', 'Có ý nghĩa', 'Không ý nghĩa sau Bonferroni'],
        ['CSES Ứng phó vấn đề', '0,17 *', '0,17–0,26 * DUY TRÌ'],
        ['CSES Tìm kiếm hỗ trợ XH', 'Có ý nghĩa', 'Có ý nghĩa DUY TRÌ'],
        ['Kết nối trường học', 'Không ý nghĩa', 'Không ý nghĩa'],
    ],
    widths=[6.0, 4.5, 5.0])
P(doc, '* có ý nghĩa thống kê sau hiệu chỉnh Bonferroni', italic=True, size=10)

P(doc, 'Phát hiện quan trọng: tự hiệu quả ứng phó tập trung vào vấn đề và tìm kiếm hỗ trợ xã '
       'hội DUY TRÌ ý nghĩa thống kê ở cả 2 và 6 tháng — cho thấy kỹ năng được học có thể được '
       'giữ lại lâu dài, dù triệu chứng trầm cảm quay lại.', bold=True)

H(doc, 'THẢO LUẬN', 2)
P(doc, 'Bối cảnh hiệu lực: Cohen d = 0,11 cho trầm cảm nằm ở persentile thứ 25 so với các '
       'chương trình phòng ngừa phổ quát khác, nhưng vượt qua 29 % các chương trình phòng ngừa '
       'SKTT cho sinh viên đại học. Meta-analysis cho thấy kích thước hiệu ứng trung vị là '
       '0,18 cho phòng ngừa trầm cảm. Vậy Happy House ở mức "trung bình–dưới" so với các chương '
       'trình phổ quát quốc tế.')

P(doc, 'Vấn đề bền vững: Khác với các thử nghiệm RAP tại Úc và New Zealand cho thấy hiệu quả '
       'duy trì 10–18 tháng, nghiên cứu này và các nghiên cứu LMIC khác cho thấy fade-out ở 6 '
       'tháng. Các tác giả cho rằng phần nguyên nhân là do phong tỏa COVID-19 tại Việt Nam '
       'trong thời gian theo dõi gây ra sự xấu đi rộng rãi của SKTT.')

P(doc, 'Tỷ lệ trầm cảm ban đầu: 25,5 % HS có CESD-R ≥ 16 ban đầu ở cả hai nhóm — nhất quán '
       'với dữ liệu Việt Nam trước đó (23 % trong các nghiên cứu tương tự) và ước tính toàn cầu.')

P(doc, 'Kỹ năng ứng phó: Các cải thiện ở cả ba chiến lược ứng phó (giải quyết vấn đề, điều '
       'tiết cảm xúc, tìm kiếm hỗ trợ xã hội) được duy trì ở 6 tháng, phù hợp với phát hiện '
       'từ thử nghiệm RAP Mauritius. Điều này gợi ý rằng Happy House phát triển KỸ NĂNG có '
       'thể giữ lại lâu dài, dù hiệu ứng trực tiếp trên TRIỆU CHỨNG giảm dần.')

P(doc, 'Thích ứng cỡ nhóm: RAP gốc được cung cấp cho ~16 HS/lớp; Happy House được thích ứng '
       'cho 40–45 HS/lớp để cải thiện khả năng triển khai/mở rộng trong bối cảnh nguồn lực '
       'hạn chế. Cần chia nhóm nhỏ hơn trong một số hoạt động và có trợ giảng hỗ trợ bổ sung.')

H(doc, 'HẠN CHẾ', 2)
P(doc, '(1) Phân bổ lớp không ngẫu nhiên (thiết kế cụm ngăn việc ngẫu nhiên hoá trong trường); '
       '(2) Đo kết cục tự báo cáo (khả năng bias phản hồi); (3) CESD-R là sàng lọc, không phải '
       'chẩn đoán; (4) Fidelity tự đánh giá bởi giáo viên, không có quan sát chất lượng bên '
       'ngoài (COVID ngăn cản nhóm Úc quan sát); (5) Đại dịch COVID-19 xảy ra trong thời gian '
       'theo dõi; (6) Phân tích trung gian (mediation analyses) không được thực hiện do hiệu '
       'ứng không có ý nghĩa ở 6 tháng.')

H(doc, 'KẾT LUẬN VÀ Ý NGHĨA', 2)
P(doc, 'Chương trình Happy House cho thấy tiềm năng lớn để tích hợp vào chương trình học hiện '
       'có nhằm giảm các vấn đề sức khoẻ tâm thần ở vị thành niên Việt Nam. Các tác giả đề '
       'xuất: (1) bổ sung BOOSTER SESSIONS để duy trì hiệu quả; (2) phân tích kinh tế cho các '
       'quyết định chính sách; (3) giải quyết các rào cản thương mại hoá (RAP yêu cầu phí bản '
       'quyền để triển khai).')

H(doc, 'QUAN ĐIỂM PHẢN BIỆN', 2)
Pred(doc, 'Điểm mạnh:', bold=True)
Pred(doc, '• ĐÂY LÀ RCT CLUSTER ĐẦU TIÊN tại Việt Nam về can thiệp SKTT học đường cho VTN — '
         'lấp khoảng trống nghiên cứu mà v2 của báo cáo can thiệp của chúng ta đã nhận định sai '
         'là "0 RCT VN".')
Pred(doc, '• Cỡ mẫu rất lớn (n = 1.084) với tỷ lệ tham gia 96,1 % — thiết kế rất chắc chắn.')
Pred(doc, '• Mất mẫu cực kỳ thấp (< 3 % ở mọi thời điểm).')
Pred(doc, '• RAP-A có cơ sở bằng chứng quốc tế mạnh (Úc, New Zealand, Mauritius) + được thích '
         'ứng văn hoá Việt Nam một cách hệ thống qua 4 giai đoạn.')
Pred(doc, '• Mô hình hợp tác Việt – Úc (Monash + Hanoi University of Public Health + QUT) — '
         'bài học quý cho hợp tác quốc tế.')
Pred(doc, '• Giáo viên cung cấp — phù hợp bối cảnh LMIC thiếu chuyên gia tâm lý.')
Pred(doc, '• Đã đăng ký thử nghiệm trước ACTRN12620000088943 — minh bạch phương pháp.')
Pred(doc, '• Tài trợ quốc tế uy tín: NHMRC Úc + NAFOSTED Việt Nam.')

Pred(doc, 'Hạn chế:', bold=True)
Pred(doc, '• Đây là can thiệp PHỔ QUÁT (universal), không phải TARGETED cho HS có triệu chứng. '
         'Hiệu ứng nhỏ d = 0,11 có thể do dilution — phù hợp với cảnh báo của Zhang, Liang & '
         'Kang 2026 (Bayesian MA 31 RCT).')
Pred(doc, '• Hiệu ứng trên trầm cảm KHÔNG CÒN ý nghĩa ở 6 tháng — FADE-OUT — gợi ý cần BOOSTER '
         'SESSIONS (các thử nghiệm RAP Úc duy trì 10–18 tháng).')
Pred(doc, '• Đo TRẦM CẢM (CESD-R) chính, KHÔNG đo LO ÂU riêng — báo cáo VTN VN cho lo âu vẫn '
         'chưa có RCT chuyên biệt.')
Pred(doc, '• Chỉ Hà Nội + lớp 10 — chưa đại diện toàn VN (lớp 11-12 chuẩn bị Gao Kao có thể '
         'khác, THCS khác, vùng nông thôn + DTTS khác).')
Pred(doc, '• Phân bổ NHÓM không ngẫu nhiên (non-randomized) — không phải RCT thuần. Có thể '
         'có bias về đặc điểm không quan sát được giữa 4 trường can thiệp vs 4 trường đối chứng.')
Pred(doc, '• KHÔNG có active control — chỉ là chương trình học chính quy, khó loại trừ hiệu '
         'ứng chú ý.')
Pred(doc, '• Fidelity tự đánh giá bởi giáo viên — không có quan sát độc lập.')
Pred(doc, '• Theo dõi trong thời gian COVID-19 — có thể confound tác động của can thiệp.')

Pred(doc, 'Áp dụng cho đề cương Việt Nam:', bold=True)
Pred(doc, '(1) Mô hình RAP-A thích ứng văn hoá VN đã KHẢ THI — tiết kiệm chi phí phát triển '
         'nội dung mới, có thể xây dựng tiếp.')
Pred(doc, '(2) Cần BOOSTER SESSIONS ở 3, 6 tháng để duy trì hiệu quả — giải quyết vấn đề '
         '"fade-out" mà Happy House gặp phải.')
Pred(doc, '(3) Nên thử nghiệm mô hình TARGETED (HS có CESD-R ≥ 16 hoặc GAD-7 ≥ 8) thay vì '
         'universal — có thể có hiệu ứng lớn hơn.')
Pred(doc, '(4) Đo đồng thời LO ÂU (GAD-7, DASS-Y, SIAS-17) + trầm cảm + coping self-efficacy '
         '— bao phủ cả giảm triệu chứng và tăng nguồn lực bảo vệ.')
Pred(doc, '(5) Mở rộng địa bàn ra ngoài Hà Nội (TPHCM, Đà Nẵng, vùng miền trung nông thôn) '
         'để tăng tính đại diện.')
Pred(doc, '(6) RANDOMIZE cấp trường để có thiết kế RCT thuần.')

P(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. Cluster controlled trial ĐẦU TIÊN tại VN cho VTN, mẫu lớn, '
       'tác giả Monash–QUT–HUPH uy tín, tỷ lệ tham gia 96,1 %, mất mẫu < 3 %. Hạn chế: '
       'universal interventions với hiệu ứng nhỏ, không đo lo âu riêng, phân bổ non-randomized. '
       'PHẢI ĐƯỢC TRÍCH DẪN trong mọi báo cáo can thiệp VN từ nay trở đi.', bold=True)

out_path = os.path.join(DICH_DIR, 'VN030_Tran_HappyHouse_Cambridge_2023.docx')
doc.save(out_path)
d = Document(out_path)
chars = sum(len(p.text) for p in d.paragraphs)
print(f'Bản dịch đầy đủ saved: {os.path.basename(out_path)}')
print(f'Chars: {chars}, Tables: {len(d.tables)}')
