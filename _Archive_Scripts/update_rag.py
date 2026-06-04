# -*- coding: utf-8 -*-
"""Tạo RAG database cho 35 bài NC + cross-study"""
import os, sys, json
os.environ['PYTHONIOENCODING'] = 'utf-8'

import chromadb
import docx

BASE = os.path.dirname(os.path.abspath(__file__))
RAG_DIR = os.path.join(BASE, 'rag_db')
DICH_DIR = os.path.join(BASE, '03_Ban-dich')
TT_DIR = os.path.join(BASE, 'Tom-tat-tung-bai')

client = chromadb.PersistentClient(path=RAG_DIR)

# Delete old collection if exists
try:
    client.delete_collection('lo_au_papers')
except:
    pass

col = client.create_collection(
    name='lo_au_papers',
    metadata={'description': 'Lo âu ở HS THCS/THPT — 35 bài NC, bản dịch + tóm tắt'}
)

def read_docx(filepath):
    """Read all text from a DOCX file"""
    try:
        doc = docx.Document(filepath)
        text = '\n'.join([p.text for p in doc.paragraphs if p.text.strip()])
        # Also read tables
        for table in doc.tables:
            for row in table.rows:
                text += '\n' + ' | '.join([c.text.strip() for c in row.cells])
        return text
    except:
        return ''

# ===== Build entries =====
entries = []

# Paper metadata
papers = [
    {'id': 'VN01', 'file_dich': '02_Hoa_2024_Frontiers.docx', 'file_tt': 'VN1_Hoa_2024.docx',
     'title': 'Lo âu HS THPT Hà Nội sau COVID-19', 'authors': 'Hoa et al. 2024', 'journal': 'Frontiers in Public Health Q1',
     'n': 3910, 'location': 'Hà Nội, VN', 'tool': 'GAD-7', 'key_finding': 'Lo âu 40,6%, nữ>nam'},
    {'id': 'VN02', 'file_dich': '06_VNAMHS_2022.docx', 'file_tt': 'VN2_VNAMHS_2022.docx',
     'title': 'V-NAMHS — Khảo sát quốc gia SKTT VTN VN', 'authors': 'V-NAMHS 2022', 'journal': 'Báo cáo quốc gia',
     'n': 5996, 'location': '38 tỉnh VN', 'tool': 'DISC-5/DSM-5', 'key_finding': 'Chẩn đoán 2,3%, tiếp cận 8,4%'},
    {'id': 'VN03', 'file_dich': '10_Pham_2024_VN_SocialSupport.docx', 'file_tt': 'VN3_Pham_2024.docx',
     'title': 'Hỗ trợ xã hội và SKTT VTN SSF VN', 'authors': 'Pham et al. 2024', 'journal': 'J Affective Disorders',
     'n': 500, 'location': 'Huế, VN', 'tool': 'DASS-21', 'key_finding': 'Hỗ trợ xã hội bảo vệ lo âu'},
    {'id': 'VN14', 'file_dich': '14_HoangTrungHoc_2025_AJPR.docx', 'file_tt': 'VN14_HoangTrungHoc_2025.docx',
     'title': 'SKTT VTN VN trước và sau COVID-19', 'authors': 'Hoàng Trung Học 2025', 'journal': 'AJPR',
     'n': 2000, 'location': 'VN', 'tool': 'DASS-21', 'key_finding': 'Lo âu 41,5%→25,4% (phục hồi)'},
    {'id': 'VN15', 'file_dich': '15_NgoAnhVinh_2024_JAffectDisordReports.docx', 'file_tt': 'VN15_NgoAnhVinh_2024.docx',
     'title': 'SKTT VTN DTTS Lạng Sơn', 'authors': 'Ngô Anh Vinh et al. 2024', 'journal': 'J Affective Disorders Reports',
     'n': 845, 'location': 'Lạng Sơn, VN (DTTS)', 'tool': 'DASS-21 + ACEs', 'key_finding': 'Lo âu 54,4%, ACEs liên quan'},
    {'id': 'VN16', 'file_dich': '16_BaoQuyen_2025_YHCD.docx', 'file_tt': 'VN16_BaoQuyen_2025.docx',
     'title': 'SKTT HS THPT Hà Nội', 'authors': 'Bảo Quyên et al. 2025', 'journal': 'TC Y học Cộng đồng',
     'n': 501, 'location': 'Hà Nội, VN', 'tool': 'DASS-21', 'key_finding': 'Lo âu 86,2% — CAO NHẤT'},
    {'id': 'VN17', 'file_dich': '17_NguyenDanhLam_2022_YHVN.docx', 'file_tt': 'VN17_DanhLam_2022.docx',
     'title': 'Stress, lo âu, trầm cảm HS THPT Thanh Hóa', 'authors': 'Nguyễn Danh Lâm 2022', 'journal': 'TC Y học VN',
     'n': 482, 'location': 'Thanh Hóa, VN', 'tool': 'DASS-21', 'key_finding': 'Lo âu 49%, tự hại 10%'},
    {'id': 'VN18', 'file_dich': '18_AnGiang_2025_YHVN.docx', 'file_tt': 'VN18_AnGiang_2025.docx',
     'title': 'Sàng lọc DASS-21 HS An Giang', 'authors': 'An Giang 2025', 'journal': 'TC Y học VN',
     'n': 366, 'location': 'An Giang, VN', 'tool': 'DASS-21', 'key_finding': 'Lo âu 61,2%'},
    {'id': 'VN19', 'file_dich': '19_TranThaoVi_2025_TLH.docx', 'file_tt': 'VN19_TranThaoVi_2025.docx',
     'title': 'Lo âu, trầm cảm và lạc quan ở VTN Huế', 'authors': 'Thảo Vi 2025', 'journal': 'TC Tâm lý học',
     'n': 685, 'location': 'Huế, VN', 'tool': 'DASS-21 + LOT-R', 'key_finding': 'Lo âu 65,8%, lạc quan trung gian'},
    {'id': 'VN20', 'file_dich': '20_TranHoVinhLoc_2024_YHTPHCM.docx', 'file_tt': 'VN20_TranHoVinhLoc_2024.docx',
     'title': 'DAS và yếu tố liên quan HS THPT TPHCM', 'authors': 'Vĩnh Lộc et al. 2024', 'journal': 'TC Y học TPHCM',
     'n': 976, 'location': 'TPHCM, VN', 'tool': 'DASS-Y', 'key_finding': 'Lo âu 25,1% (DASS-Y thấp hơn DASS-21)'},
    {'id': 'QT21', 'file_dich': '21_Norway_2025_SocSciMed.docx', 'file_tt': 'QT21_Norway_2025.docx',
     'title': 'Giải thích xu hướng tăng distress VTN Na Uy 2011-2024', 'authors': 'Brunborg et al. 2025', 'journal': 'Social Science & Medicine Q1',
     'n': 979043, 'location': 'Na Uy', 'tool': 'HSCL-6', 'key_finding': 'Bất mãn trường = giải thích chính; MXH một phần'},
    {'id': 'QT22', 'file_dich': '22_ScreenTime_2025_BJCP.docx', 'file_tt': 'QT22_ScreenTime_2025.docx',
     'title': 'Screen time và trầm cảm/lo âu VTN — dọc 12 tháng', 'authors': 'Li et al. 2025', 'journal': 'BJCP Q1',
     'n': 4058, 'location': 'Úc', 'tool': 'PHQ-A + CAS-8', 'key_finding': 'Cắt ngang mạnh, DỌC YẾU; lo âu dọc p=0,443'},
    {'id': 'QT23', 'file_dich': '23_JAACAP_US_Trends_2024.docx', 'file_tt': 'QT23_JAACAP_US_2024.docx',
     'title': 'Xu hướng rối loạn tâm thần trẻ em Mỹ 2013-2021', 'authors': 'Mojtabai & Olfson 2024', 'journal': 'JAACAP Q1 IF=11',
     'n': 13684154, 'location': 'Mỹ', 'tool': 'Chẩn đoán lâm sàng MH-CLD', 'key_finding': 'Lo âu TĂNG GẤP ĐÔI (AOR=2,17)'},
    {'id': 'QT24', 'file_dich': '24_WHO_Europe_2025_LancetRegional.docx', 'file_tt': 'QT24_WHO_Europe_2025.docx',
     'title': 'SKTT trẻ em/thanh niên khu vực WHO châu Âu', 'authors': 'Tarasenko et al. 2025', 'journal': 'Lancet Regional Health Europe Q1 IF=15',
     'n': 0, 'location': 'Châu Âu (53 nước)', 'tool': 'Tổng quan chính sách', 'key_finding': '9 triệu VTN có RLSKTT, lo âu+trầm cảm >50%'},
    {'id': 'QT25', 'file_dich': '25_EpiPsychSci_2025.docx', 'file_tt': 'QT25_EpiPsychSci_2025.docx',
     'title': 'CMH toàn diện VTN Úc 2018-2022', 'authors': 'Crisp et al. 2025', 'journal': 'EpiPsychSci Q1 IF=7',
     'n': 5656, 'location': 'Úc', 'tool': 'K10 + MHC-SF', 'key_finding': 'Flourishing giảm 53%→44,4%; khoảng cách giới mở rộng'},
    {'id': 'QT26', 'file_dich': '26_UK_NHS_2025_Parliament.docx', 'file_tt': 'QT26_UK_NHS_2025.docx',
     'title': 'Thống kê SKTT Anh: tỷ lệ, dịch vụ, chi tiêu', 'authors': 'Baker & Kirk-Wade 2024', 'journal': 'UK Parliament',
     'n': 0, 'location': 'England', 'tool': 'NHS Digital', 'key_finding': 'Lo âu 7-16: 3,5%→6,3%; £16,3 tỷ/năm'},
    {'id': 'QT27', 'file_dich': '27_NatureHumanBehav_SocialMedia_2025.docx', 'file_tt': 'QT27_Nature_SocialMedia_2025.docx',
     'title': 'MXH ở VTN có/không SKTT — Registered Report', 'authors': 'Fassi et al. 2025', 'journal': 'Nature Human Behaviour Q1 IF=30',
     'n': 3340, 'location': 'UK', 'tool': 'DAWBA (chẩn đoán)', 'key_finding': 'VTN SKTT dùng MXH nhiều hơn (g=0,46); nội hóa > ngoại hóa'},
    {'id': 'QT28', 'file_dich': '28_AJP_PediatricAnxiety_2024.docx', 'file_tt': 'QT28_AJP_Treatment_2024.docx',
     'title': 'Điều trị lo âu trẻ em: hiện tại và tương lai', 'authors': 'Zugman et al. 2024', 'journal': 'AJP Q1 IF=18',
     'n': 0, 'location': 'Review', 'tool': 'Tổng quan', 'key_finding': 'CBT+SSRI 80,7% đáp ứng (CAMS); NNT=4'},
    {'id': 'QT29', 'file_dich': '29_CBT_NetworkMeta_2025_BMCPsych.docx', 'file_tt': 'QT29_CBT_NetworkMeta_2025.docx',
     'title': 'NMA can thiệp lo âu trẻ em: ACT, CBT, PE, VRET', 'authors': 'Li et al. 2025', 'journal': 'BMC Psychiatry Q1',
     'n': 1711, 'location': '12 nước', 'tool': 'NMA Bayesian (30 RCT)', 'key_finding': 'ACT hạng 1 (SUCRA 0,69); CBT hạng 2 (bằng chứng mạnh nhất)'},
    {'id': 'QT30', 'file_dich': '30_GBD_Trends_10-24y_2025.docx', 'file_tt': 'QT30_GBD_Trends.docx',
     'title': 'Xu hướng trầm cảm/lo âu 10-24 tuổi 1990-2021 (GBD)', 'authors': 'Zhang et al. 2025', 'journal': 'J Affective Disorders Q1',
     'n': 0, 'location': '204 nước', 'tool': 'GBD 2021 Joinpoint', 'key_finding': 'Lo âu AAPC 0,84%/năm; 10-14 tuổi tăng nhanh nhất'},
    {'id': 'QT31', 'file_dich': '31_59Countries_Anxiety_2025.docx', 'file_tt': 'QT31_59Countries.docx',
     'title': 'Lo âu VTN đi học 59 quốc gia (GSHS)', 'authors': 'Islam et al. 2025', 'journal': 'J Affective Disorders Q1',
     'n': 179937, 'location': '59 nước LMIC', 'tool': 'GSHS (WHO)', 'key_finding': 'Bất an thực phẩm AOR=2,22; ý tưởng tự tử AOR=2,84'},
    {'id': 'QT32', 'file_dich': '32_Ireland_MyWorld_2024.docx', 'file_tt': 'QT32_Ireland.docx',
     'title': 'Xu hướng trầm cảm/lo âu VTN Ireland 2012-2019', 'authors': 'Fitzgerald et al. 2024', 'journal': 'Early Intervention in Psychiatry',
     'n': 11954, 'location': 'Ireland', 'tool': 'DASS-21', 'key_finding': 'Tăng trước COVID; nữ nhanh hơn; OGA bảo vệ'},
    {'id': 'QT33', 'file_dich': '33_JAMA_ScreenMedia_2024.docx', 'file_tt': 'QT33_JAMA_Screen.docx',
     'title': 'RCT giảm screen time → cải thiện SKTT trẻ em', 'authors': 'Schmidt-Persson et al. 2024', 'journal': 'JAMA Network Open Q1 IF=13,8',
     'n': 181, 'location': 'Đan Mạch', 'tool': 'SDQ (RCT)', 'key_finding': 'Cohen d=0,53; nội hóa cải thiện mạnh nhất'},
    {'id': 'QT34', 'file_dich': '34_Korea_MH_Trends_2024.docx', 'file_tt': 'QT34_Korea_Trends.docx',
     'title': 'Xu hướng SKTT VTN Hàn Quốc theo thu nhập 2006-2022', 'authors': 'Cho et al. 2024', 'journal': 'Nature Sci Rep Q1',
     'n': 0, 'location': 'Hàn Quốc', 'tool': 'KYRBS', 'key_finding': 'Đảo chiều: giảm trước COVID→tăng sau; nghèo 62,8% vs giàu 40,1%'},
    {'id': 'QT35', 'file_dich': '35_SocialAnxiety_7Countries_2020.docx', 'file_tt': 'QT35_SocialAnxiety_7Countries.docx',
     'title': 'Lo âu xã hội ở thanh niên 7 quốc gia (bao gồm VN)', 'authors': 'Jefferies & Ungar 2020', 'journal': 'PLOS ONE Q1',
     'n': 6825, 'location': '7 nước (VN, TQ, ID, RU, TH, US, BR)', 'tool': 'SIAS-17', 'key_finding': 'VN SAD=30,7%; tổng 36%; không khác biệt giới; 18% false negatives'},
]

print(f'Processing {len(papers)} papers...')

documents = []
metadatas = []
ids = []

for p in papers:
    # Read summary
    tt_path = os.path.join(TT_DIR, p['file_tt'])
    tt_text = read_docx(tt_path) if os.path.exists(tt_path) else ''

    # Read translation (first 3000 chars for RAG)
    dich_path = os.path.join(DICH_DIR, p['file_dich'])
    dich_text = read_docx(dich_path)[:3000] if os.path.exists(dich_path) else ''

    # Combine for RAG
    combined = f"""BÀI {p['id']}: {p['title']}
Tác giả: {p['authors']}
Tạp chí: {p['journal']}
Mẫu: n={p['n']:,} | Địa bàn: {p['location']}
Công cụ: {p['tool']}
Phát hiện chính: {p['key_finding']}

--- TÓM TẮT ---
{tt_text[:2000]}

--- BẢN DỊCH (trích) ---
{dich_text[:1500]}"""

    documents.append(combined)
    metadatas.append({
        'paper_id': p['id'],
        'title': p['title'],
        'authors': p['authors'],
        'journal': p['journal'],
        'n': str(p['n']),
        'location': p['location'],
        'tool': p['tool'],
        'key_finding': p['key_finding'],
    })
    ids.append(p['id'])

# Add cross-study synthesis
cs_path = os.path.join(BASE, 'Tổng hợp liên bài báo - Lo âu HS - 04042026.docx')
if os.path.exists(cs_path):
    cs_text = read_docx(cs_path)
    documents.append(f"TỔNG HỢP LIÊN BÀI BÁO — 35 NC\n{cs_text}")
    metadatas.append({'paper_id': 'CROSS_STUDY', 'title': 'Cross-Study Synthesis', 'authors': 'Tổng hợp', 'journal': '', 'n': '35', 'location': 'Toàn cầu', 'tool': 'Liên bài', 'key_finding': '9 bảng tổng hợp'})
    ids.append('CROSS_STUDY')

# Add to ChromaDB
col.add(documents=documents, metadatas=metadatas, ids=ids)
print(f'Added {len(documents)} entries to RAG')
print(f'Collection count: {col.count()}')

# Test query
results = col.query(query_texts=['tỷ lệ lo âu ở học sinh Việt Nam'], n_results=3)
print('\nTest query: "tỷ lệ lo âu ở học sinh Việt Nam"')
for i, (doc, meta) in enumerate(zip(results['documents'][0], results['metadatas'][0])):
    print(f'  {i+1}. [{meta["paper_id"]}] {meta["title"]} — {meta["key_finding"][:60]}')

results2 = col.query(query_texts=['can thiệp CBT cho lo âu trẻ em'], n_results=3)
print('\nTest query: "can thiệp CBT cho lo âu trẻ em"')
for i, (doc, meta) in enumerate(zip(results2['documents'][0], results2['metadatas'][0])):
    print(f'  {i+1}. [{meta["paper_id"]}] {meta["title"]} — {meta["key_finding"][:60]}')

print('\nDONE!')
