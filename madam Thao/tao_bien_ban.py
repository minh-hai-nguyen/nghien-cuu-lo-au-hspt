# -*- coding: utf-8 -*-
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def make_doc(include_stories=True):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)

    def add_h(text, level=1):
        h = doc.add_heading(text, level=level)
        for r in h.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0,0,0)

    def add_p(text, bold=False, italic=False):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.bold = bold
        r.italic = italic
        return p

    def shade(cell, color):
        s = OxmlElement('w:shd')
        s.set(qn('w:fill'), color)
        s.set(qn('w:val'), 'clear')
        cell._tc.get_or_add_tcPr().append(s)

    suffix = ' (Ban day du)' if include_stories else ' (Ban tom tat)'
    title = doc.add_heading('', level=0)
    r = title.add_run('BIEN BAN CUOC HOP' + suffix + '\nChuong trinh dao tao lanh dao \u2014 Hoi nhap & Toa sang')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 51, 102)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    add_p('Chu tri: Madam Thao', bold=True)
    add_p('Thanh phan: Tuan va cac thanh vien nhom dao tao, truyen thong')
    add_p('Noi dung: Xay dung chuong trinh dao tao lanh dao ket hop hoi nhap toa sang cho Vietjet')
    doc.add_paragraph()

    add_h('I. Tinh than cot loi cua chuong trinh dao tao', 1)
    add_p('Madam Thao nhan manh rang chuong trinh dao tao phai di vao chieu sau, khong chi dua tren kich ban co san. Muc tieu la tao moi truong de nguoi tham gia tim thay gia tri ban than trong tap the. Tinh than ban dau cua cac chuong trinh dao tao tai Vietjet da bi mai mot \u2014 thieu yeu to con nguoi, thieu "hon".')
    add_p('Chi muon tro lai chuong trinh dao tao voi diem khac biet nam o yeu to con nguoi: khi gap kho khan, tap trung vao cach vuot qua va phoi hop doi nhom.')

    add_h('II. Thuc trang nhan su', 1)
    add_p('Madam Thao chia se thang than: tinh than nhan vien Vietjet nhung nam qua da giam sut dang ke. Nguoi ta roi di khong phai vi che do (che do tot), ma vi moi truong \u2014 khong tim thay y nghia cong viec.')
    add_p('Hinh anh ngay xua: to bay briefing tai tru so cong ty, xep hang ngay ngan, dieu trong van phong \u2014 nhan vien nhin thay to bay chuan bi lam nhiem vu thi tu hao. Bay gio: dau tat mat toi, ra khoi nha 2-3 gio sang, ra san bay roi ve di ngu, khong co thoi gian tieu, khong thay y nghia gan bo.')

    add_h('III. Ke hoach chuong trinh cu the', 1)
    add_p('Quy mo: 60 nguoi, chia 5 team', bold=True)
    add_p('Boi canh: Chuong trinh long ghep voi khoa hoc 6 thang "Toa sang Tuong lai". Lam 2 trong 1: 1 tieng dau danh cho tinh than Vietjet, phan con lai la hoi nhap toa sang danh cho lanh dao.')

    add_p('Ngay 1:', bold=True)
    add_p('\u2022 Buoi sang: Khoi dong voi truyen thong Vietjet, truyen lua (~1 tieng); tiep theo hoi nhap toa sang cho lanh dao')
    add_p('\u2022 Noi dung lanh dao: cong tac lap ke hoach, quan ly tai chinh, quan ly nhan su, ngan sach, chi tieu, chi so')
    add_p('\u2022 Buoi chieu: Noi dung chuyen sau + gioi thieu chinh thuc')
    add_p('\u2022 Toi: Gala Dinner, team building (tap luyen tu toi hom truoc)')

    add_p('Ngay 2:', bold=True)
    add_p('\u2022 Dua vao cac noi dung ky nang')
    add_p('\u2022 Tham khao chuong trinh doi tac (A11, INSEAD va cac doi tac khac)')

    add_p('Team building:', bold=True)
    add_p('\u2022 5 team, moi team co team truong. Tap luyen tu toi thu 6')
    add_p('\u2022 Game trong hoi truong. Thiet ke dong phuc team cho thi dau')
    add_p('\u2022 Kinh nghiem tu Nha Trang va Galaxy: tinh than rat tot, khong ngai')

    add_h('IV. Phan cong cong viec', 1)
    t = doc.add_table(rows=5, cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(['Nguoi', 'Nhiem vu']):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shade(c, 'D9E2F3')
    for ri, (k, v) in enumerate([
        ('Tuan', 'Giai quyet viec gia dinh, bao cao lai cho nhom'),
        ('Nhom dao tao', 'Ghi nhan noi dung, gui lai tai lieu da lam'),
        ('Chu Tuan', 'Chuan bi Agenda theo huong dan cua chi'),
        ('Madam Thao', 'Chot kich ban, chat loc noi dung tu tai lieu doi tac'),
    ]):
        t.rows[ri+1].cells[0].text = k
        t.rows[ri+1].cells[1].text = v
        for ci in range(2):
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    doc.add_paragraph()

    add_h('V. Triet ly lanh dao \u2014 Madam Thao', 1)
    add_p('Nhan qua: Lam tot thi se co ket qua tot. Khong can xu ly cang thang bang quyen luc. Moi nguoi co quyen sai va duoc tao co hoi khac phuc.')
    add_p('Cong bang: Can thoi gian va co che, khong phai quyet dinh cam tinh. Neu lanh dao cap tren khong guong mau ma ky luat cap duoi thi khong cong bang.')
    add_p('Coaching: Moi lanh dao phai tro thanh nguoi coach cho cap duoi. Tinh than hoc tap phai duoc duy tri lien tuc.')

    if include_stories:
        add_h('VI. Cau chuyen truyen cam hung \u2014 Madam Thao ke', 1)

        add_h('1. Nang cong suat san bay Tan Son Nhat', 2)
        add_p('"Nam 2013 minh lam, 2014 bat dau chay. Du an nang cao nang luc dieu hanh cat ha canh o san bay Tan Son Nhat. Nang len gap doi tu 30 luot den 60 luot. Bay gio dang van hanh 48 luot. Tu 24 trieu luot khach len 42 trieu luot khach. Du an ay lam cho ca nganh hang khong. ACV muoi may nam lien, moi nam don 16 trieu luot khach, loi nhuan may tram trieu do la/nam."')

        add_h('2. Giai cuu san chung khoan HoSE', 2)
        add_p('"Giai cuu cho ca thi truong von. Ca nuoc My co san New York, Nasdaq. Ca nuoc Anh co London Stock Exchange. Ma no bi treo thi giai cuu may tram ty do la o day. Hon 100 cong ty chung khoan, 4 trieu khach hang to chuc va ca nhan trong nuoc va ngoai nuoc. Moi ngay giao dich 1-2 ty do la. Minh cung nhay vao giai cuu. The con minh thi khong can cai gi ca."')

        add_h('3. Ngoai giao qua mua may bay', 2)
        add_p('"Mua may bay Boeing la di voi nuoc My \u2014 ong Trump, Obama, Bush. Mua may bay Airbus la di voi Phap, Duc, Tay Ban Nha. Mua COMAC cua Trung Quoc la chuong trinh ong Tap Can Binh. Tao nen xay dung moi quan he ngoai giao, xay dung kinh te giua cac quoc gia, dan toc. Giam bot di chien tranh, xung dot."')

        add_h('4. APEC va nang luc trien khai than toc', 2)
        add_p('"APEC 2 lan, nam 2017. Dua Viet Nam vao the gioi va dua the gioi ve Viet Nam. Giai phap truong sang minh lam co 90 ngay \u2014 Han Quoc mat 17 nam. Cung hoi nghi xay 4 thang."')

        add_h('5. Thien nguyen COVID va triet ly thiet ke', 2)
        add_p('"Mua Covid, lao di cho vac-xin, thiet bi y te, bao gao. Thiet ke nha chung cu phai co thang may lon cho cang cuu thuong va quan tai. O dau co bac thang phai co duong doc cho xe noi, xe nguoi khuyet tat. Day la tinh than minh dua vao moi hoat dong."')

        add_h('6. Tam voc toan dien', 2)
        add_p('"Co vua len top 7 the gioi. Bong da nam vao World Cup, bong da nu cung vao World Cup. Toan dien cho dat nuoc, o tat ca cac linh vuc \u2014 van hoa, xa hoi, kinh te, tai chinh."')

        add_p('\u2192 Madam Thao nhan manh: "Moi con nguoi deu co the tro nen vi dai. Lam nhung viec lon co y nghia, tu nhung hoat dong don gian, gian di."', italic=True)

    step = 'VII' if include_stories else 'VI'
    add_h(f'{step}. Ket luan va buoc tiep theo', 1)
    add_p('1. Chot kich ban chuong trinh 2 ngay')
    add_p('2. Chuan bi tai lieu dao tao \u2014 tham khao tu doi tac, chat loc voi chi')
    add_p('3. Long ghep chuong trinh hoi nhap toa sang voi su kien Vietjet')
    add_p('4. Phan chia 5 team, chuan bi team building tu toi hom truoc')
    add_p('5. Gui lai tai lieu va noi dung da lam cho nhom')
    add_p('6. Xay dung clip ve Madam Thao de phuc vu chuong trinh')
    add_p('7. Viet lai cau chuyen Madam Thao \u2014 can nguoi biet ke chuyen, co tu lieu')

    doc.add_paragraph()
    add_p('Bien ban tao tu ban ghi am qua Whisper large-v3 (faster-whisper, GPU RTX 3060). Da doi chieu 3 ban transcript (small + medium + large) va hieu dinh.', italic=True)

    return doc

doc1 = make_doc(include_stories=True)
doc1.save('Bien ban hop - Day du - Madam Thao.docx')
print('Done: Bien ban hop - Day du - Madam Thao.docx')

doc2 = make_doc(include_stories=False)
doc2.save('Bien ban hop - Tom tat - Madam Thao.docx')
print('Done: Bien ban hop - Tom tat - Madam Thao.docx')
