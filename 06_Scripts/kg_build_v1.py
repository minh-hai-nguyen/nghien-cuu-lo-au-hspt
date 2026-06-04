# -*- coding: utf-8 -*-
"""
KNOWLEDGE GRAPH v1 — Du an Lo au.
Schema:
  Paper --AUTHORED_BY--> Author
  Paper --PUBLISHED_IN--> Journal
  Paper --CONDUCTED_IN--> Country
  Paper --USED_METHOD--> Method
  Paper --MEASURED--> Outcome
  Paper --USED_SCALE--> Scale
  Paper --REPORTED_ES--> EffectSize
  Paper --HAS_N--> SampleSize
  Paper --HAS_YEAR--> Year

Sources:
  1. canonical_mapping.json (60 papers)
  2. canonical_index.json (file paths)
  3. Tom-tat-tung-bai/*.docx (structured summaries — extract via tables)
  4. 03_Ban-dich/*.docx (for extra details)

Output:
  - kg_data/nodes.json
  - kg_data/edges.json
  - kg_data/graph.graphml (for Gephi)
  - kg_data/graph.html (interactive pyvis)
  - kg_data/validation_report.md
"""
import os, sys, io, json, re
from collections import defaultdict, Counter
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import networkx as nx
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
KG_DIR = os.path.join(os.path.dirname(__file__), 'kg_data')
os.makedirs(KG_DIR, exist_ok=True)

# ============================================================
# 1. LOAD CANONICAL MAPPING + INDEX
# ============================================================
with open(os.path.join(os.path.dirname(__file__), 'canonical_mapping.json'), encoding='utf-8') as f:
    CANONICAL = json.load(f)

with open(os.path.join(ROOT, '02_Papers-goc', 'canonical_index.json'), encoding='utf-8') as f:
    INDEX = json.load(f)

print(f'Loaded {len(CANONICAL)} canonical, {len(INDEX)} indexed')

# ============================================================
# 2. EXTRACT STRUCTURED DATA FROM SUMMARIES
# ============================================================
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')

def extract_summary(canon_id):
    """Read summary docx and extract structured fields from info table."""
    # Find file
    file = None
    for f in os.listdir(TT_DIR):
        if f.startswith(canon_id + '_') and f.endswith('.docx'):
            file = f
            break
    if not file:
        return None
    path = os.path.join(TT_DIR, file)
    d = Document(path)
    # Extract all tables
    fields = {'file': file, 'id': canon_id}
    # Look for table with "Tác giả", "Tạp chí", "Năm", etc.
    for tb in d.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2:
                key = cells[0]
                val = ' | '.join(cells[1:])
                if key:
                    fields[key] = val
    # Also extract full text for later
    all_text = '\n'.join(p.text for p in d.paragraphs if p.text.strip())
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text
    fields['_full_text'] = all_text
    return fields

# Extract all summaries
summaries = {}
for item in CANONICAL:
    canon = item['canonical']
    canon3 = canon[:2] + f'{int(canon[2:]):03d}'
    s = extract_summary(canon3)
    if s:
        summaries[canon3] = s
        s['descriptor'] = item['descriptor']

print(f'\nExtracted {len(summaries)} summary entities')

# ============================================================
# 3. PARSE FIELDS INTO TRIPLES
# ============================================================
G = nx.MultiDiGraph()
triples = []  # (subj, pred, obj)

# Helper to normalize entity names
def normalize(s):
    return re.sub(r'\s+', ' ', s).strip()

# Parse each summary
for canon_id, data in summaries.items():
    # Add paper node
    desc = data.get('descriptor', canon_id)
    G.add_node(canon_id, type='Paper', descriptor=desc, file=data.get('file', ''))

    # Extract authors
    authors_field = data.get('Tác giả', '') or data.get('Tác giả:', '')
    if authors_field:
        # Split by common delimiters
        authors = re.split(r'[,;]|\band\b|\&|và cộng sự', authors_field)
        for i, a in enumerate(authors[:5]):  # max 5 authors
            a = normalize(a)
            a = re.sub(r'^\s*\d+[,\-–]?\s*', '', a)  # remove leading numbers
            a = a.replace('et al.', '').strip()
            if len(a) >= 4 and not any(c.isdigit() for c in a[:5]):
                # Add author node
                author_id = f'Author::{a[:40]}'
                if not G.has_node(author_id):
                    G.add_node(author_id, type='Author', name=a)
                G.add_edge(canon_id, author_id, rel='AUTHORED_BY')
                triples.append((canon_id, 'AUTHORED_BY', author_id))

    # Extract journal
    journal_field = data.get('Tạp chí', '') or data.get('Journal', '')
    if journal_field:
        # Extract journal name (before first comma/paren)
        j = re.split(r'[(,]', journal_field)[0].strip()
        if j:
            jid = f'Journal::{j[:40]}'
            if not G.has_node(jid):
                G.add_node(jid, type='Journal', name=j)
            G.add_edge(canon_id, jid, rel='PUBLISHED_IN')
            triples.append((canon_id, 'PUBLISHED_IN', jid))

    # Extract year
    year_match = re.search(r'(20\d{2})', data.get('Xuất bản', '') or data.get('Năm', '') or data.get('_full_text', '')[:500])
    if year_match:
        year = year_match.group(1)
        yid = f'Year::{year}'
        if not G.has_node(yid):
            G.add_node(yid, type='Year', value=int(year))
        G.add_edge(canon_id, yid, rel='HAS_YEAR')
        triples.append((canon_id, 'HAS_YEAR', yid))

    # Extract n (sample size)
    n_match = re.search(r'[Nn]\s*=\s*(\d+[\.,]?\d*)', data.get('Mẫu', '') or data.get('_full_text', ''))
    if n_match:
        n_str = n_match.group(1).replace('.', '').replace(',', '')
        try:
            n = int(n_str)
            nid = f'SampleSize::{n}'
            if not G.has_node(nid):
                G.add_node(nid, type='SampleSize', value=n)
            G.add_edge(canon_id, nid, rel='HAS_N')
            triples.append((canon_id, 'HAS_N', nid))
        except: pass

    # Extract country/region
    country_hints = {
        'Việt Nam': 'Vietnam', 'Vietnam': 'Vietnam', 'VN': 'Vietnam',
        'Hà Nội': 'Vietnam', 'TPHCM': 'Vietnam', 'Huế': 'Vietnam', 'Hải Phòng': 'Vietnam',
        'Long An': 'Vietnam', 'Vĩnh Long': 'Vietnam', 'An Giang': 'Vietnam',
        'Trung Quốc': 'China', 'China': 'China', 'TQ': 'China',
        'Hàn Quốc': 'Korea', 'Korea': 'Korea',
        'Nhật': 'Japan', 'Japan': 'Japan',
        'Sri Lanka': 'Sri Lanka',
        'Ấn Độ': 'India', 'India': 'India',
        'UK': 'UK', 'Anh': 'UK', 'United Kingdom': 'UK',
        'Mỹ': 'USA', 'USA': 'USA', 'United States': 'USA',
        'Úc': 'Australia', 'Australia': 'Australia',
        'Philippines': 'Philippines', 'Indonesia': 'Indonesia',
        'ASEAN': 'ASEAN', 'ĐNA': 'Southeast Asia', 'Đông Nam Á': 'Southeast Asia',
        'LMIC': 'LMIC',
        'Ethiopia': 'Ethiopia',
        'Ireland': 'Ireland', 'Norway': 'Norway', 'Đức': 'Germany',
        'Denmark': 'Denmark', 'Đan Mạch': 'Denmark',
    }
    full_text = data.get('_full_text', '')
    for keyword, country in country_hints.items():
        if keyword in full_text[:2000]:  # first 2000 chars
            cid = f'Country::{country}'
            if not G.has_node(cid):
                G.add_node(cid, type='Country', name=country)
            if not G.has_edge(canon_id, cid):
                G.add_edge(canon_id, cid, rel='CONDUCTED_IN')
                triples.append((canon_id, 'CONDUCTED_IN', cid))

    # Extract methods
    method_hints = {
        'CBT': 'CBT', 'Cognitive Behavioral': 'CBT',
        'iCBT': 'iCBT', 'Internet CBT': 'iCBT', 'internet-based CBT': 'iCBT',
        'gCBT': 'gCBT', 'CBT nhóm': 'gCBT',
        'Mobile CBT': 'Mobile CBT', 'app CBT': 'Mobile CBT',
        'Mindfulness': 'Mindfulness',
        'DMHI': 'DMHI',
        'resilience': 'Resilience training',
        'Resilience': 'Resilience training',
        'CA-CBT': 'CA-CBT',
        'Physical exercise': 'Physical Exercise',
        'hoạt động thể chất': 'Physical Exercise',
        'Yoga': 'Yoga', 'Thư giãn': 'Relaxation',
        'VR': 'VR Therapy', 'virtual reality': 'VR Therapy',
        'SSRI': 'SSRI', 'Sertraline': 'SSRI',
    }
    for keyword, method in method_hints.items():
        if keyword in full_text:
            mid = f'Method::{method}'
            if not G.has_node(mid):
                G.add_node(mid, type='Method', name=method)
            if not G.has_edge(canon_id, mid):
                G.add_edge(canon_id, mid, rel='USED_METHOD')
                triples.append((canon_id, 'USED_METHOD', mid))

    # Extract outcomes
    outcome_hints = {
        'lo âu xã hội': 'Social Anxiety (SAD)',
        'Social Anxiety': 'Social Anxiety (SAD)', 'SAD': 'Social Anxiety (SAD)',
        'GAD': 'Generalized Anxiety (GAD)',
        'lo âu tổng quát': 'Generalized Anxiety (GAD)',
        'lo âu': 'Anxiety',
        'Anxiety': 'Anxiety',
        'trầm cảm': 'Depression', 'Depression': 'Depression',
        'căng thẳng': 'Stress', 'Stress': 'Stress',
    }
    for keyword, outcome in outcome_hints.items():
        if keyword in full_text:
            oid = f'Outcome::{outcome}'
            if not G.has_node(oid):
                G.add_node(oid, type='Outcome', name=outcome)
            if not G.has_edge(canon_id, oid):
                G.add_edge(canon_id, oid, rel='MEASURED')
                triples.append((canon_id, 'MEASURED', oid))

    # Extract scales
    scale_hints = ['DASS-21', 'DASS-Y', 'DASS-42', 'GAD-7', 'HAM-A', 'PHQ-9', 'SIAS', 'SCAS',
                   'CES-D', 'STAI', 'LSAS', 'ASI', 'PSQI', 'LOT-R', 'ESSA', 'YBRS']
    for scale in scale_hints:
        if scale in full_text:
            sid = f'Scale::{scale}'
            if not G.has_node(sid):
                G.add_node(sid, type='Scale', name=scale)
            if not G.has_edge(canon_id, sid):
                G.add_edge(canon_id, sid, rel='USED_SCALE')
                triples.append((canon_id, 'USED_SCALE', sid))

    # Extract effect sizes (Cohen d, Hedges g, OR, SMD, β, SUCRA)
    es_pats = [
        (r'[Cc]ohen[\'\u2019]?s?\s*d\s*=\s*(-?\d+[,.]\d+)', 'Cohen_d'),
        (r'[Hh]edges[\'\u2019]?s?\s*g\s*=\s*(-?\d+[,.]\d+)', 'Hedges_g'),
        (r'\bd\s*=\s*(-?\d+[,.]\d+)', 'd'),
        (r'\bg\s*=\s*(-?\d+[,.]\d+)', 'g'),
        (r'\bOR\s*=\s*(-?\d+[,.]\d+)', 'OR'),
        (r'\bAOR\s*=\s*(-?\d+[,.]\d+)', 'AOR'),
        (r'\bSMD\s*=\s*(-?\d+[,.]\d+)', 'SMD'),
        (r'β\s*=\s*(-?\d+[,.]\d+)', 'beta'),
        (r'\bSUCRA\s*[=:]\s*(\d+[,.]\d+|\d+\s*%)', 'SUCRA'),
    ]
    for pat, es_type in es_pats:
        for m in re.finditer(pat, full_text):
            val_str = m.group(1).replace(',', '.').replace('%', '').strip()
            try:
                val = float(val_str)
                es_node = f'ES::{canon_id}::{es_type}::{val}'
                if not G.has_node(es_node):
                    G.add_node(es_node, type='EffectSize', es_type=es_type, value=val, paper=canon_id)
                G.add_edge(canon_id, es_node, rel='REPORTED_ES')
                triples.append((canon_id, 'REPORTED_ES', es_node))
            except: pass

# ============================================================
# 4. SAVE GRAPH
# ============================================================
print(f'\n=== GRAPH STATS ===')
print(f'Nodes: {G.number_of_nodes()}')
print(f'Edges: {G.number_of_edges()}')
node_types = Counter(G.nodes[n].get('type', 'Unknown') for n in G.nodes)
print(f'Node types: {dict(node_types)}')

edge_rels = Counter(d['rel'] for _, _, d in G.edges(data=True))
print(f'Edge types: {dict(edge_rels)}')

# Save
nx.write_graphml(G, os.path.join(KG_DIR, 'kg_v1.graphml'))
print(f'\nSaved GraphML: {os.path.join(KG_DIR, "kg_v1.graphml")}')

# Save triples as JSON
with open(os.path.join(KG_DIR, 'kg_triples.json'), 'w', encoding='utf-8') as f:
    json.dump([{'subject': s, 'predicate': p, 'object': o} for s, p, o in triples],
              f, ensure_ascii=False, indent=2)
print(f'Saved triples: {len(triples)}')

# Save nodes + edges
nodes_data = [{'id': n, **G.nodes[n]} for n in G.nodes]
edges_data = [{'source': u, 'target': v, **d} for u, v, d in G.edges(data=True)]
with open(os.path.join(KG_DIR, 'nodes.json'), 'w', encoding='utf-8') as f:
    json.dump(nodes_data, f, ensure_ascii=False, indent=2)
with open(os.path.join(KG_DIR, 'edges.json'), 'w', encoding='utf-8') as f:
    json.dump(edges_data, f, ensure_ascii=False, indent=2)

print('\n=== TOP PAPERS BY DEGREE ===')
for n in sorted(G.nodes(), key=lambda x: G.degree(x), reverse=True)[:15]:
    if G.nodes[n].get('type') == 'Paper':
        print(f'  {n}: degree={G.degree(n)} | {G.nodes[n].get("descriptor", "")}')
