# -*- coding: utf-8 -*-
"""
Add the 9 missing images to VN022 translation.
Insert each after its corresponding (or nearest preceding) page marker.
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from copy import deepcopy

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/unicef_imgs/'

# 9 missing images — tell me which page marker to insert after, plus caption
# Based on original PDF content context (UNICEF Viet Nam 2022 report)
MISSING = [
    # (filename, anchor_page, caption_vn, caption_en)
    ('p006_img1_Im0.jp2', 6,
     'Hình 2: Ảnh minh hoạ đầu chương Lời nói đầu',
     'Figure 2: Foreword chapter opener illustration (p.6)'),
    ('p011_img1_Im0.jpg', 11,
     'Hình 3: Ảnh minh hoạ đầu Chương 1 — Giới thiệu',
     'Figure 3: Chapter 1 opener — Introduction (p.11)'),
    ('p023_img1_Im0.jpg', 22,  # no p23 marker, use p22
     'Hình 4: Ảnh minh hoạ đầu Chương 2 — Phương pháp nghiên cứu',
     'Figure 4: Chapter 2 opener — Methodology (p.23)'),
    ('p028_img1_Im0.jpg', 27,  # no p28 marker, use p27
     'Hình 5: Ảnh minh hoạ — Bối cảnh trường học tại Việt Nam',
     'Figure 5: School context in Viet Nam (p.28)'),
    ('p035_img1_Im0.jpg', 34,  # no p35 marker, use p34
     'Hình 7: Ảnh minh hoạ — Sức khoẻ tâm thần học sinh',
     'Figure 7: Student mental health illustration (p.35)'),
    ('p043_img1_Im0.jpg', 42,  # no p43 marker, use p42
     'Hình 9: Ảnh minh hoạ đầu chương — Yếu tố trường học',
     'Figure 9: School factors chapter opener (p.43)'),
    ('p055_img1_Im0.jpg', 54,  # no p55 marker, use p54
     'Hình 11: Ảnh minh hoạ — Mối quan hệ giáo viên–học sinh',
     'Figure 11: Teacher–student relationships (p.55)'),
    ('p066_img1_Im0.jpg', 65,  # no p66 marker, use p65
     'Hình 13: Ảnh minh hoạ — Bắt nạt và bắt nạt mạng',
     'Figure 13: Bullying and cyberbullying (p.66)'),
    ('p075_img1_Im0.jpg', 74,  # no p75 marker, use p74
     'Hình 15: Ảnh minh hoạ đầu chương Khuyến nghị',
     'Figure 15: Recommendations chapter opener (p.75)'),
]

d = Document(PATH)

# Find all page marker paragraphs and their XML elements
markers = {}
for i, p in enumerate(d.paragraphs):
    m = re.match(r'---\s*Trang\s+(\d+)', p.text.strip())
    if m:
        page = int(m.group(1))
        if page not in markers:
            markers[page] = p

# Insert an image + caption paragraph right AFTER a given paragraph
def insert_image_after(anchor_para, img_path, caption_vn, caption_en, width_cm=11.0):
    # Create image paragraph
    img_p = d.add_paragraph()
    img_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        img_p.add_run().add_picture(img_path, width=Cm(width_cm))
    except Exception as e:
        print(f'    FAILED to add picture: {e}')
        # Remove empty paragraph
        img_p._element.getparent().remove(img_p._element)
        return False

    # Caption paragraph
    cap_p = d.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap_p.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Now move both paragraphs (appended at end of doc) to right after anchor
    img_el = img_p._element
    cap_el = cap_p._element
    anchor_el = anchor_para._element

    # Remove from current (end) position
    img_el.getparent().remove(img_el)
    cap_el.getparent().remove(cap_el)

    # Insert after anchor: put cap first after anchor, then img (so img appears first)
    anchor_el.addnext(cap_el)
    anchor_el.addnext(img_el)
    return True

added = 0
for fname, page, cvn, cen in MISSING:
    img_path = os.path.join(IMG_DIR, fname)
    if not os.path.exists(img_path):
        print(f'  SKIP {fname}: file not found')
        continue
    if page not in markers:
        print(f'  SKIP {fname}: page marker {page} not found')
        continue
    anchor = markers[page]
    ok = insert_image_after(anchor, img_path, cvn, cen)
    if ok:
        added += 1
        print(f'  OK: {fname} inserted after "Trang {page}"')

d.save(PATH)
print(f'\nTotal added: {added}/{len(MISSING)}')

# Verify
import zipfile
with zipfile.ZipFile(PATH) as z:
    media = [n for n in z.namelist() if n.startswith('word/media/')]
    print(f'Images now in docx: {len(media)}')
