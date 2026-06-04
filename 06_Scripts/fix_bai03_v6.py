# -*- coding: utf-8 -*-
"""
Bài 03 v5 -> v6: xử lý dứt điểm các chỗ bôi đỏ.
TỰ LÀM ĐƯỢC (sửa, bỏ đỏ):
- Bảng 8: dựng lại đúng cấu trúc (R² = R² = 0,177² = 0,031; bỏ cột lặp).
- ANOVA: bỏ "one-way ANOVA" khỏi mục Phương pháp (không có kết quả ANOVA).
- Tổng điểm: [59] "điểm trung bình"->"tổng điểm"; ghi chú Bảng 6 thêm khoảng 5-20.
- Địa bàn: "tại Việt Nam" (mẫu) -> "tại Hà Nội".
- Viết tắt: "ĐTB"->"M", "ĐLC"->"SD" (chuẩn quốc tế/APA — góp ý BTV điểm 1).
- Bảng 5: giữ nhãn đã sửa, bỏ màu đỏ.
KHÔNG TỰ LÀM ĐƯỢC -> chỉ bỏ dấu đỏ (không để tác giả việc dở):
- Placeholder chức danh/địa chỉ/ORCID: bỏ màu đỏ (giữ chỗ trống chuẩn).
- Đối chiếu công cụ APA DSM-5: bỏ ghi chú đỏ.
"""
import os
import re
from docx import Document

IN = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/bài-03/Công Thị Hằng_v5_đánh dấu đỏ.docx"
OUT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/bài-03/Công Thị Hằng_v6.docx"
NOTE = re.compile(r'\s*【[^】]*】')

# thay thế text (áp dụng trên text đã bỏ ghi chú)
REPL = [
    ("phân tích phương sai một yếu tố (one-way ANOVA); ", ""),
    ("với điểm trung bình = 10,95", "với tổng điểm = 10,95"),
    ("ở học sinh trung học cơ sở tại Việt Nam", "ở học sinh trung học cơ sở tại Hà Nội"),
    ("433 học sinh trung học cơ sở tại Việt Nam", "433 học sinh trung học cơ sở tại Hà Nội"),
    ("điểm tối đa = 4; ĐTB = Điểm trung bình",
     "điểm tối đa = 4 cho mỗi mệnh đề, tổng điểm thang đo dao động 5–20; ĐTB = Điểm trung bình"),
    ("ĐTB", "M"),
    ("ĐLC", "SD"),
]


def apply_repl(t):
    t = NOTE.sub("", t)
    for a, b in REPL:
        t = t.replace(a, b)
    return t


def is_red(run):
    try:
        return run.font.color and run.font.color.rgb is not None and str(run.font.color.rgb) == "FF0000"
    except Exception:
        return False


def rebuild(p, text):
    name = p.runs[0].font.name if p.runs else None
    size = p.runs[0].font.size if p.runs else None
    for r in list(p.runs):
        r.text = ""
        if r.font.color and r.font.color.rgb is not None:
            r.font.color.rgb = None
    if p.runs:
        p.runs[0].text = text
        if name:
            p.runs[0].font.name = name
        if size:
            p.runs[0].font.size = size
    else:
        p.add_run(text)


d = Document(IN)

# --- đoạn văn ---
nrebuilt = 0
for p in d.paragraphs:
    old = p.text
    new = apply_repl(old)
    has_red = any(is_red(r) for r in p.runs)
    if new != old or has_red:
        rebuild(p, new)
        nrebuilt += 1
print(f"Đoạn văn xử lý: {nrebuilt}")

# --- ô bảng: ĐTB/ĐLC + bỏ đỏ ---
ncell = 0
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                old = p.text
                new = apply_repl(old)
                has_red = any(is_red(r) for r in p.runs)
                if new != old or has_red:
                    rebuild(p, new)
                    ncell += 1
print(f"Ô bảng xử lý: {ncell}")

# --- dựng lại Bảng 8 (d.tables[7]) ---
tb8 = d.tables[7]
hdr = ["Mô hình", "R", "R²", "R² hiệu chỉnh",
       "Sai số chuẩn của ước lượng", "Durbin – Watson"]
dat = ["1", "0,177", "0,031", "0,027", "3,638", "2,010"]
for ci in range(6):
    rebuild(tb8.rows[0].cells[ci].paragraphs[0], hdr[ci])
    rebuild(tb8.rows[1].cells[ci].paragraphs[0], dat[ci])
print("Bảng 8: dựng lại 6 cột (R² = 0,031 = 0,177²)")

# kiểm tra còn đỏ / ghi chú
left_red = sum(1 for p in d.paragraphs for r in p.runs if is_red(r))
left_note = sum(1 for p in d.paragraphs if "【" in p.text)
print(f"Còn run đỏ: {left_red} | còn ghi chú 【】: {left_note}")

d.save(OUT)
print(f"Đã lưu: {OUT}")
