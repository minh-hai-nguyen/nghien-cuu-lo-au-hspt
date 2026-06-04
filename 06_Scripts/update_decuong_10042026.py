# -*- coding: utf-8 -*-
"""
Update đề cương + báo cáo (05042026 → 10042026)
- Mở doc cũ
- Thêm PHẦN IV: Cập nhật 14 bài mới (45-58)
- Cập nhật TLTK
"""
import sys, os, shutil, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Lo au - Bao cao + De cuong - 05042026.docx')
DST = os.path.join(ROOT, 'Lo au - Bao cao + De cuong - 10042026.docx')

shutil.copy(SRC, DST)
doc = Document(DST)

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    tcW.append(w)

def H(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def P(text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb = color

def table(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
    return t

doc.add_page_break()
H('PHẦN IV — CẬP NHẬT 10/04/2026 (14 bài bổ sung — tổng 58 bài)', level=1)
P('Cập nhật 10/04/2026: Bổ sung 14 bài mới (10 PDF đầy đủ + 4 abstract-only paywall) — bài 45-58. Phần này tổng hợp các điều chỉnh CẦN ÁP DỤNG cho đề cương đã có ở Phần III.', italic=True)

H('IV.1. CẬP NHẬT TỔNG QUAN TÀI LIỆU (cho mục 1)', level=2)
P('Bổ sung vào đặt vấn đề các bằng chứng mới sau 10/04/2026:', bold=True)
P('• Tỷ lệ lo âu HS VN đã có thêm dữ liệu: Hải Phòng THPT 39,3% DASS-21 (#45), Long An THPT 57,2% (#47, CAO NHẤT VN). Bản đồ tỷ lệ lo âu VN dao động 16% (SV GAD-7≥10) → 57,2% (Long An), phản ánh khác biệt vùng + công cụ.')
P('• Bằng chứng meta lớn nhất hiện có cho COVID era: Chen 2025 (#55) — 141 NC + 1 TRIỆU trẻ — xác định 14 yếu tố BẢO VỆ và 29 yếu tố NGUY CƠ. "Emotional functioning" + family support + community resources là buffer chính.')
P('• Phát hiện CƠ CHẾ mới: Dong 2025 PLOS (#54, n=2.716 TQ) — KÊNH GIAO TIẾP gia đình giảm 78% nguy cơ trầm cảm (OR=0,22), tâm sự bạn giảm 63% (OR=0,37). Quan trọng hơn cả mức độ "hỗ trợ" chung. Điều này GỢI Ý hướng can thiệp mới: tập huấn KỸ NĂNG GIAO TIẾP cha-con.')
P('• Cơ chế trung gian: Zheng 2025 (#46) — MXH ảnh hưởng lo âu CHỦ YẾU qua self-efficacy (gián tiếp 63%), không trực tiếp. Giấc ngủ là YẾU TỐ MẠNH NHẤT (β=0,615).')
P('• Bằng chứng CBT trường MÂU THUẪN: Zhang 2026 Bayesian MA 31 RCT n=19.865 (#56) — CBT phổ quát hiệu quả NHỎ, chất lượng nền THẤP. Nhưng B11 Japan iCBT subthreshold SAD RCT (#51) DƯƠNG TÍNH. → CBT TARGETED tốt hơn UNIVERSAL.')
P('• Mobile CBT: Qiaochu 2025 SR 9 RCT n=2.479 (#57) — mạnh cho TRẦM CẢM (7/8 NC) nhưng YẾU cho LO ÂU (2/6). B3 JAMA App CBT 2024 (#49) cung cấp RCT mới. Khi cho SAD-specific (B2, B11) thì có hiệu quả.')
P('• Khoảng trống KHU VỰC: Menon 2025 Scoping (#58) — 69 NC / 12 nước LMIC Đông Á + TBD — xác định khoảng trống ở can thiệp CỘNG ĐỒNG, GIA ĐÌNH, dài hạn. VN nằm trong khoảng trống này.')

H('IV.2. CẬP NHẬT MỤC TIÊU NGHIÊN CỨU (cho mục 2)', level=2)
P('Bổ sung mục tiêu cụ thể (5):', bold=True)
P('(5) Đánh giá vai trò của KÊNH GIAO TIẾP gia đình (frequency, content, confidant) như yếu tố bảo vệ chính — so sánh với mức độ "hỗ trợ" chung. Phù hợp Dong 2025 (#54) và Chen 2025 (#55).')

H('IV.3. CẬP NHẬT PHƯƠNG PHÁP — Giai đoạn 1 (Khảo sát)', level=2)
P('Bổ sung công cụ + biến số đo:', bold=True)
P('• Thêm KÊNH GIAO TIẾP gia đình: hỏi (a) tần suất tâm sự cha/mẹ; (b) chủ đề tâm sự; (c) người tin tưởng nhất (family/peer/specific person/no confidant) — như Dong 2025 (#54).')
P('• Thêm SELF-EFFICACY (GSES) — biến trung gian quan trọng (Zheng 2025, #46).')
P('• Đo CẢ 3 trục DASS-21 (trầm cảm + lo âu + stress) — không chỉ lo âu — để có dữ liệu so sánh trực tiếp với Dong 2025 (#54), Hải Phòng (#45), Long An (#47).')
P('• Thêm câu hỏi "FEAR OF LETTING DOWN OTHERS" (sợ làm người khác thất vọng) — Dong 2025 phát hiện 60,3% HS TQ trải qua. Có thể là yếu tố văn hóa Á Đông.')
P('• Phân tích đa biến: dùng cả Model 1 (đơn biến) và Model 2 (kiểm soát confounders với VIF/p>0,1) — như Dong 2025.')

H('IV.4. CẬP NHẬT PHƯƠNG PHÁP — Giai đoạn 2 (RCT can thiệp)', level=2)
P('THIẾT KẾ CAN THIỆP CẬP NHẬT:', bold=True)
P('• KHÔNG dùng mô hình UNIVERSAL phổ quát (bằng chứng yếu — Zhang 2026 #56, B5 mindfulness UK 8.376 HS thất bại).')
P('• Dùng mô hình TARGETED + SELF-REFERRAL: HS có triệu chứng tự đăng ký tham gia (PLACES/BESST UK — B5). Ngôn ngữ thường ngày ("căng thẳng" thay "trầm cảm") để giảm kỳ thị.')
P('• THÀNH PHẦN can thiệp 12 tuần (cập nhật):')
P('  (1) CBT NHÓM TARGETED — như cũ, nhưng targeted (HS GAD-7 ≥8 hoặc DASS-Y lo âu ≥8). BMC NMA QT29 + B11 Japan iCBT.')
P('  (2) MODULE GIAO TIẾP CHA-CON (mới — phát hiện Dong A5): 4 buổi, kỹ năng "tâm sự" — listening, không phán xét, đề xuất câu hỏi. Cha mẹ + HS cùng tham gia. Phù hợp Menon B10 (gap can thiệp gia đình).')
P('  (3) HOẠT ĐỘNG THỂ CHẤT (PE) — như cũ, BMC NMA QT29 PE SUCRA 0,51.')
P('  (4) MODULE RESILIENCE (mới — B6): 3 buổi, dựa trên SR+MA RCTs B6 (#50). Thành phần: kỹ năng ứng phó, lạc quan, kết nối xã hội, tự nhận thức.')
P('  (5) MOBILE CBT TIẾNG VIỆT (mới — B3 JAMA + B4 Qiaochu): tối thiểu 1 module trầm cảm (bằng chứng mạnh) + 1 module lo âu (yếu hơn nhưng có B11 Japan SAD-specific).')
P('  (6) BỎ "OGA mentor" hoặc giữ song song (Ireland QT32) — vẫn có bằng chứng.')
P('• ĐO ĐỦ trước-sau-3 tháng-12 tháng (B1 Bayesian MA #56 cho thấy hiệu quả duy trì 1 năm).')

H('IV.5. CẬP NHẬT Ý NGHĨA KHOA HỌC', level=2)
P('Bổ sung điểm mới của đề tài:', bold=True)
P('• ĐẦU TIÊN tại VN đo KÊNH GIAO TIẾP gia đình như yếu tố bảo vệ độc lập — kiểm chứng phát hiện Dong 2025 (#54) trên dân số VN.')
P('• ĐẦU TIÊN tại VN thử nghiệm CAN THIỆP PHỐI HỢP CBT-targeted + module giao tiếp gia đình + resilience + mobile CBT — lấp khoảng trống Menon (#58) về can thiệp đa cấp ở LMIC Đông Á.')
P('• ĐẦU TIÊN tại VN dùng PLACES/self-referral model (B5 UK) — đánh giá chấp nhận văn hóa.')
P('• So sánh universal CBT vs targeted CBT — đối chiếu với Zhang 2026 Bayesian (#56).')

H('IV.6. CẬP NHẬT KHOẢNG TRỐNG NC (Top 17)', level=2)
table(
    ['#', 'Khoảng trống', 'Bằng chứng', 'Ưu tiên'],
    [
        ['11', 'RCT can thiệp KỸ NĂNG GIAO TIẾP gia đình ở VN', 'Dong A5 OR=0,22; Menon B10', '★★★★★'],
        ['12', 'CBT TARGETED self-referral cho VTN VN có triệu chứng', 'B5 BESST UK; B11 Japan; Zhang B1 (universal yếu)', '★★★★'],
        ['13', 'Mobile CBT tiếng Việt — TRẦM CẢM mạnh + SAD-specific', 'B3 JAMA; B4 Qiaochu (87,5% trầm cảm); B11 Japan', '★★★★'],
        ['14', 'NC dọc đo CẢ 3 trục DASS-21 mẫu VN >2.000 THPT', 'Dong A5; VN21 Trần Thảo Vi', '★★★'],
        ['15', 'Can thiệp CỘNG ĐỒNG + GIA ĐÌNH ở LMIC Đông Á (gồm VN)', 'Menon B10 scoping', '★★★'],
        ['16', 'Chuẩn hóa công cụ đo lo âu vùng VN (16% → 57%)', 'Hoa, Hải Phòng, Long An, VN COVID, Dinh', '★★★'],
        ['17', 'Resilience trường tích hợp CBT — RCT VN', 'B6 MA + Ireland QT32 + VN21 lạc quan', '★★'],
    ],
    widths=[1.0, 6.5, 6.0, 1.5]
)

H('IV.7. TÀI LIỆU THAM KHẢO BỔ SUNG (14 bài 45-58)', level=2)
P('Cập nhật danh mục TLTK chính:', bold=True)
P('Chen, H., Wang, Q., Zhu, J., et al. (2025). Protective and risk factors of anxiety in children and adolescents during COVID-19: A systematic review and three level meta-analysis. J Affective Disorders, 374, 408–432.')
P('Dong, T., Wang, Y., & Lin, Y. (2025). Prevalence and determinants of depression, anxiety, and stress among secondary school students. PLOS ONE, 20(9), e0328785.')
P('Zheng, G.F., & Peng, H.Y. (2025). The effects of social media addiction, academic stress, and sleep quality on anxiety symptoms: A cross-sectional study of Chinese vocational students. Psychology Research and Behavior Management, 18, 1571-1584.')
P('Zhang, X., Liang, Z., & Kang, J. (2026). Long-term effects of school-based CBT in low-risk children and adolescents: A Bayesian meta-analysis. Journal of Clinical Psychology, 82, 248–259.')
P('Zhang, Q., & Wang, Y. (2025). Effectiveness of mobile-based CBT for depression and anxiety in children and young people: A SR of RCTs. Clinical Psychology & Psychotherapy, 32(6), e70173.')
P('Menon, V., Coppard, M., McEwen, S., et al. (2025). Evaluated interventions targeting mental health in children and adolescents: Scoping review LMIC East Asia and Pacific. Asia Pacific J Public Health, 37(4), 332–346.')
P('Brown, J.S.L., & Carter, B. (2025). School-based interventions for depression and anxiety in UK. Journal of Mental Health, 34(4), 357–361.')
P('Cao, C., et al. (2025). School-based interventions for resilience in children and adolescents: SR + MA of RCTs. Frontiers in Psychiatry, 16, 1594658.')
P('Praptomojati, A., et al. (2024). Culturally-adapted CBT for anxiety in Southeast Asian populations: A systematic review of 7 studies. [B7]')
P('De Silva, S., et al. (2024). Cluster RCT of school-based CBT in Sri Lanka (n=720). [B8]')
P('Xian, X., et al. (2024). Network meta-analysis of interventions for social anxiety disorder in adolescents (30 RCTs, n=1.547). [B9]')
P('Walder, N., et al. (2025). Digital mental health interventions for social anxiety disorder: A meta-analysis (21 RCTs). JMIR. [B2]')
P('Nguyen, L.X., et al. (2023). Anxiety in Vietnamese university students during COVID-19 (n=5.730, GAD-7). Medicine. [A8]')
P('UNICEF Vietnam (2022). School factors causing Vietnamese adolescents anxiety. [A3]')
P('Trần Thảo Vi, et al. (2024). Longitudinal 3-year study on anxiety in junior high school students in Hue. J Rural Med. [A1]')

doc.save(DST)
print(f'Saved: {DST}')

d2 = Document(DST)
print(f'Total paragraphs: {len(d2.paragraphs)}, tables: {len(d2.tables)}')
