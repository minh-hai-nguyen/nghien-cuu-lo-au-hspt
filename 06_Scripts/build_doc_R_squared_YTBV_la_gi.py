"""DOC: Tra loi cau hoi cua thay 'R² YTBV goi la gi?'
Don gian, de hieu cho thay khong chuyen thong ke.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/R_squared_YTBV_la_gi.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('R² YTBV GỌI LÀ GÌ?\n— Trả lời câu hỏi của thầy 09/05/2026 —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue('R² YTBV gọi là gì Em ơi?', italic=True)

# Trả lời ngắn
H('Trả lời NGẮN GỌN', level=2, color=NAVY)
para(
    'R² gọi là HỆ SỐ XÁC ĐỊNH (tiếng Anh: Coefficient of '
    'Determination; còn gọi: hệ số quyết định, hệ số giải '
    'thích phương sai). R² YTBV = TỶ LỆ PHƯƠNG SAI rối loạn '
    'lo âu được GIẢI THÍCH bởi yếu tố bảo vệ tổng.', bold=True
)
para('Cụ thể trong chương 3 luận án:', bold=True)
para('     R² YTBV = 0,124  →  Yếu tố bảo vệ tổng giải thích 12,4% phương sai RLLA')

# I. Tên gọi
H('I. Tên gọi đầy đủ', level=2, color=NAVY)
add_table(
    ['Ngôn ngữ', 'Tên gọi'],
    [
        ['Tiếng Anh', 'Coefficient of Determination (viết tắt R² hoặc R-squared)'],
        ['Tiếng Việt CHUẨN', 'Hệ số xác định'],
        ['Tiếng Việt phổ biến', 'Hệ số quyết định / Hệ số R²'],
        ['Tên đầy đủ', 'Tỷ lệ phương sai được giải thích (variance explained)'],
        ['Trong chương 3 luận án', 'R² (giữ ký hiệu quốc tế — phổ biến nhất)'],
    ]
)
para('')
para(
    'Lưu ý: R (chữ hoa, không có ²) là HỆ SỐ TƯƠNG QUAN BỘI '
    '(multiple correlation coefficient). R² = bình phương của R. '
    'Nên đọc R² là "R bình phương" hoặc "R-squared".', italic=True
)

# II. Ý nghĩa
H('II. Ý nghĩa của R²', level=2, color=NAVY)
para(
    'R² cho biết MỨC ĐỘ mà biến độc lập (yếu tố ảnh hưởng) GIẢI '
    'THÍCH biến phụ thuộc (kết quả).'
)
para('')
add_table(
    ['Giá trị R²', 'Ý nghĩa', 'Diễn giải'],
    [
        ['R² = 0', '0% phương sai giải thích được', 'Yếu tố HOÀN TOÀN KHÔNG ảnh hưởng'],
        ['R² = 0,5', '50% phương sai giải thích được', 'Yếu tố giải thích MỘT NỬA biến thiên'],
        ['R² = 1', '100% phương sai giải thích được', 'Yếu tố giải thích HOÀN TOÀN (lý tưởng)'],
        ['R² YTBV = 0,124', '12,4% phương sai giải thích', 'YTBV giải thích 12,4% biến thiên RLLA; còn 87,6% do các yếu tố KHÁC'],
    ]
)
para('')
para('Cách diễn giải đơn giản R² = 0,124:', bold=True)
para('• "Yếu tố bảo vệ giải thích 12,4% biến thiên rối loạn lo âu."')
para('• "Khi biết YTBV của một học sinh, ta có thể dự đoán 12,4% sự khác biệt RLLA giữa các học sinh."')
para('• "Còn lại 87,6% biến thiên RLLA là do các yếu tố KHÁC chưa đo trong mô hình này."')

# III. Quy ước đánh giá
H('III. Quy ước đánh giá độ lớn R² (Cohen 1988)', level=2, color=NAVY)
add_table(
    ['Khoảng R²', 'Đánh giá Cohen (1988)', 'Diễn giải'],
    [
        ['R² < 0,02', 'Không đáng kể', 'Yếu tố hầu như không có vai trò'],
        ['0,02 ≤ R² < 0,13', 'Nhỏ (small)', 'Tác động yếu nhưng có ý nghĩa'],
        ['0,13 ≤ R² < 0,26', 'Trung bình (medium)', 'Tác động đáng kể'],
        ['R² ≥ 0,26', 'Lớn (large)', 'Tác động mạnh, giải thích phần lớn biến thiên'],
    ]
)
para('')
para(
    'Áp dụng cho R² YTBV = 0,124: NẰM Ở RANH GIỚI giữa nhỏ '
    '(small) và trung bình (medium) — gần đạt ngưỡng "trung '
    'bình". Yếu tố bảo vệ tổng có tác động ĐÁNG KỂ nhưng KHÔNG '
    'mạnh.', bold=True
)

# IV. Áp dụng vào chương 3
H('IV. Áp dụng vào chương 3 luận án', level=2, color=NAVY)
add_table(
    ['Mô hình', 'R²', 'Đánh giá Cohen 1988', 'Bảng tham chiếu'],
    [
        ['YTNC riêng (Áp lực HT + NĐT + BNHĐ)', '0,558',
         'RẤT LỚN (vượt ngưỡng 0,26 hơn 2 lần)', 'Bảng 3.40'],
        ['YTBV riêng (TTr + GBTH + HTCM + HTBB)', '0,124',
         'Nhỏ–trung bình (gần 0,13)', 'Bảng 3.42'],
        ['Mô hình tích hợp YTNC + YTBV', '0,598',
         'RẤT LỚN', 'Bảng 3.38'],
        ['ALHT riêng → RLLA', '0,284',
         'Lớn (vượt ngưỡng 0,26)', 'Bảng 3.24'],
        ['NĐT riêng → RLLA', '0,160',
         'Trung bình', 'Bảng 3.26'],
        ['BNHĐ riêng → RLLA', '0,076',
         'Nhỏ (gần ngưỡng dưới)', 'Bảng 3.28'],
        ['BPĐP riêng → RLLA', '0,540',
         'Rất lớn (nhưng FIT KÉM)', 'Bảng 3.45 — cảnh báo'],
    ]
)
para('')
para(
    'Phát hiện đáng chú ý: R² YTBV (12,4%) THẤP HƠN nhiều R² YTNC '
    '(55,8%). Tỷ số 0,124/0,558 ≈ 0,222 — yếu tố bảo vệ tổng '
    'chỉ giải thích ~22% so với yếu tố nguy cơ tổng. Phù hợp '
    'với mô hình tâm lý học: yếu tố nguy cơ thường có cường độ '
    'mạnh hơn yếu tố bảo vệ ở giai đoạn vị thành niên.', bold=True
)

# V. Cách công thức (đơn giản)
H('V. Cách tính R² (đơn giản hóa)', level=2, color=NAVY)
para('Công thức:', bold=True)
para('     R² = (Phương sai được giải thích bởi mô hình) / (Tổng phương sai biến phụ thuộc)')
para('Tương đương:', bold=True)
para('     R² = 1 − (Phương sai phần dư) / (Tổng phương sai)')
para('')
para(
    'Hình dung trực quan: nếu vẽ scatter plot YTBV (trục x) với '
    'RLLA (trục y), R² đo MỨC ĐỘ các điểm bám vào ĐƯỜNG XU HƯỚNG '
    'của mô hình. R² càng gần 1 → các điểm càng GẦN đường; R² '
    'càng gần 0 → các điểm càng PHÂN TÁN.'
)

# VI. Cảnh báo
H('VI. Cảnh báo khi diễn giải R²', level=2, color=NAVY)
para('R² CÓ và KHÔNG có những đặc điểm sau:', bold=True)
add_table(
    ['R² CÓ thể nói gì?', 'R² KHÔNG nói được gì?'],
    [
        ['Tỷ lệ % phương sai giải thích',
         'Có quan hệ NHÂN QUẢ không'],
        ['Mô hình có "phù hợp" với data không',
         'Yếu tố cụ thể nào trong YTBV mạnh nhất'],
        ['So sánh nhiều mô hình (R² lớn hơn = giải thích tốt hơn)',
         'Hệ số β chuẩn hóa của từng đường ảnh hưởng'],
        ['Hiệu ứng tổng của các biến độc lập',
         'Mô hình SEM có "fit" hay không (cần CFI, RMSEA)'],
    ]
)
para('')
para(
    'Khi báo cáo R² trong báo cáo CTH, NÊN đi kèm: (a) β chuẩn '
    'hóa cho từng đường ảnh hưởng; (b) fit indices (CFI, RMSEA, '
    'TLI); (c) p-value của phép kiểm định F.', italic=True
)

# VII. Tóm gọn
H('VII. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue('R² YTBV gọi là HỆ SỐ XÁC ĐỊNH (Coefficient of Determination) — '
     'đo TỶ LỆ PHƯƠNG SAI rối loạn lo âu được giải thích bởi yếu '
     'tố bảo vệ tổng.', bold=True)
blue('')
blue('Trong chương 3 luận án: R² YTBV = 0,124 → "Yếu tố bảo vệ '
     'tổng giải thích 12,4% biến thiên rối loạn lo âu", còn 87,6% '
     'do các yếu tố khác.')
blue('')
blue('Đánh giá theo Cohen (1988): R² = 0,124 NẰM Ở RANH GIỚI giữa '
     'nhỏ và trung bình — yếu tố bảo vệ tổng có tác động ĐÁNG KỂ '
     'nhưng KHÔNG mạnh. Thấp hơn nhiều R² YTNC = 0,558 (rất lớn) '
     '— gợi ý ưu tiên giảm yếu tố nguy cơ trong can thiệp.')

# Phụ lục TLTK
H('Phụ lục — Tham chiếu', level=2, color=NAVY)
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.first_line_indent = Cm(-1.0)
p.add_run('Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.').font.size = Pt(11)

p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
p.paragraph_format.left_indent = Cm(1.0)
p.paragraph_format.first_line_indent = Cm(-1.0)
p.add_run('Falk, R. F., & Miller, N. B. (1992). A primer for soft modeling. University of Akron Press. [Phân loại R² 0,10/0,30/0,50 cho mô hình SEM PLS.]').font.size = Pt(11)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
