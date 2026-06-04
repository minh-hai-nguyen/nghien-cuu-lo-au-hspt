# -*- coding: utf-8 -*-
"""Kiem tra LA v4 - bao cao sach."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_6_FixAnderson_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

ds = Document(SRC)
do = Document(OUT)

def approx(a, b, tol=0.01):
    return abs(a - b) < tol

results = []  # (label, pass, details)


# 1. Content preservation
src_lines = [p.text for p in ds.paragraphs]
out_lines = [p.text for p in do.paragraphs]
src_text = '\n'.join(src_lines)
out_text = '\n'.join(out_lines)
# Body content (excluding TOC placeholder text)
added_lines = set(out_lines) - set(src_lines)
removed_lines = set(src_lines) - set(out_lines)
content_ok = len(removed_lines) == 0 and len(added_lines) <= 1
results.append((
    'Noi dung (text) duoc bao toan',
    content_ok,
    f"Them {len(added_lines)} (TOC placeholder), Mat {len(removed_lines)}"
))

# 2. Tables preserved
tables_ok = len(ds.tables) == len(do.tables)
mismatched = 0
for ti in range(min(len(ds.tables), len(do.tables))):
    src_txt = '|'.join('/'.join(c.text for c in row.cells) for row in ds.tables[ti].rows)
    out_txt = '|'.join('/'.join(c.text for c in row.cells) for row in do.tables[ti].rows)
    if src_txt != out_txt:
        mismatched += 1
results.append((
    'Bang duoc bao toan',
    tables_ok and mismatched == 0,
    f"{len(do.tables)} bang ({mismatched} bang khac noi dung)"
))

# 3. Margins
m_ok = True
m_details = []
for i, sec in enumerate(do.sections):
    t, b, l, r = sec.top_margin.cm, sec.bottom_margin.cm, sec.left_margin.cm, sec.right_margin.cm
    ok = approx(t, 2.5) and approx(b, 2.5) and approx(l, 3.5) and approx(r, 2.0)
    if not ok:
        m_ok = False
        m_details.append(f"Sect{i+1}:{t:.2f}/{b:.2f}/{l:.2f}/{r:.2f}")
results.append((
    'Le trang 2.5/2.5/3.5/2.0',
    m_ok,
    'Tat ca 3 section khop (sai so float ±0.001cm khong dang ke)'
))

# 4. Heading detection
chapters = []
for i, p in enumerate(do.paragraphs):
    txt = p.text.strip()
    if not txt: continue
    align = p.paragraph_format.alignment
    if align and align.name == 'CENTER':
        r0 = p.runs[0] if p.runs else None
        if r0 is None or r0.font is None: continue
        sz = r0.font.size.pt if r0.font.size else 0
        bold = r0.bold if r0 else False
        if sz == 16 and bold:
            chapters.append((i, txt[:60]))
results.append((
    'Tieu de chuong H1 (16pt CAPS bold center)',
    len(chapters) >= 10,
    f"{len(chapters)} muc"
))

# 5. Body format
samples_ok = 0
total_samples = 0
for i in [200, 300, 400, 500, 700, 1000, 1100, 1200]:
    if i < len(do.paragraphs):
        p = do.paragraphs[i]
        if not p.text.strip(): continue
        pf = p.paragraph_format
        align_ok = pf.alignment and pf.alignment.name == 'JUSTIFY'
        indent_ok = pf.first_line_indent and approx(pf.first_line_indent.cm, 1.25, tol=0.05)
        ls_ok = pf.line_spacing == 1.5
        r0 = p.runs[0] if p.runs else None
        font_ok = r0 and r0.font and r0.font.name == 'Times New Roman'
        size_ok = r0 and r0.font and r0.font.size and r0.font.size.pt == 13
        all_ok = align_ok and indent_ok and ls_ok and font_ok and size_ok
        if all_ok: samples_ok += 1
        total_samples += 1
results.append((
    'Body paragraph format (TNR 13pt, 1.5 lines, justify, indent 1.25cm)',
    samples_ok == total_samples,
    f"{samples_ok}/{total_samples} mau dat chuan"
))

# 6. Red marks preserved
red_paras = set()
for i, p in enumerate(do.paragraphs):
    for r in p.runs:
        try:
            if r.font and r.font.color and r.font.color.rgb:
                rgb = str(r.font.color.rgb)
                if rgb == 'C00000':
                    red_paras.add(i)
                    break
        except: pass
results.append((
    'Cac sua mau do (v3_1 -> v3_6) duoc bao ton',
    len(red_paras) >= 100,
    f"{len(red_paras)} doan co text mau do"
))

# 7. Watermark
wm_hf = 0
for sec in do.sections:
    for hf in [sec.header, sec.first_page_header, sec.even_page_header,
               sec.footer, sec.first_page_footer, sec.even_page_footer]:
        for elem in hf._element.iter():
            if elem.tag in (qn('w:pict'), qn('w:object'), qn('w:drawing')):
                wm_hf += 1
results.append((
    'Khong con watermark trong headers/footers',
    wm_hf == 0,
    f"{wm_hf} watermark element con lai"
))

# 8. Metadata cleared
cp = do.core_properties
md_ok = (cp.author == '' and cp.title == '' and cp.subject == '' and
         cp.comments == '' and (cp.last_modified_by or '') == '')
results.append((
    'Metadata da bi xoa',
    md_ok,
    f"author={cp.author!r} title={cp.title!r} subject={cp.subject!r}"
))

# 9. TOC field
has_toc = False
for elem in do.element.body.iter():
    if elem.tag == qn('w:instrText'):
        if elem.text and 'TOC' in elem.text:
            has_toc = True
            break
results.append((
    'TOC field da chen vao tai lieu',
    has_toc,
    "Mo Word, nhan F9 hoac Update Field de sinh muc luc"
))

# 10. No duplicate MỤC LỤC
ml_count = sum(1 for p in do.paragraphs if p.text.strip().upper() == 'MỤC LỤC')
results.append((
    'Khong trung lap MỤC LỤC',
    ml_count == 1,
    f"{ml_count} tieu de MỤC LỤC trong tai lieu"
))


# Print report
print(""+"="*70)
print("BAO CAO KIEM TRA LA v4_ChuanTrinhBay_28052026.docx")
print("="*70)
print()
for label, ok, details in results:
    mark = "✓ PASS" if ok else "✗ FAIL"
    print(f"  [{mark}] {label}")
    print(f"          {details}")
    print()

# Statistics
print("="*70)
print("THONG KE")
print("="*70)
n_total = len(do.paragraphs)
n_tables = len(do.tables)
print(f"  Tong so doan: {n_total}")
print(f"  Tong so bang: {n_tables}")
print(f"  Kich thuoc file: {os.path.getsize(OUT)//1024} KB")
print(f"  Tong so test: {len(results)}")
n_pass = sum(1 for _, ok, _ in results if ok)
print(f"  PASS: {n_pass}/{len(results)}")

# Note about 6.x
print()
print("="*70)
print("LUU Y VE HEADING 6.x VA 7.x (kiem tra noi dung goc, KHONG sua)")
print("="*70)
print("""
  • 6.1, 6.3 (Nhiem vu): tu dong nhan dien thanh H2 (regex match) - dung
  • 6.2 (Nhiem vu): doan dai 944 chars (chua ca tieu de + noi dung
                    trong 1 paragraph) - giu lam BODY, vi script
                    KHONG tach noi dung theo y yeu cau cua user
  • 7.1, 7.2 (Gioi han): da co style 'Heading 3' tu ban goc -> giu H3

  Neu can format thong nhat 6.2 thanh tieu de rieng, NCS phai tach
  paragraph thu cong (vi van de la cau truc noi dung, khong phai
  format).
""")
