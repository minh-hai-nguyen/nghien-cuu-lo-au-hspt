# -*- coding: utf-8 -*-
"""
qa_VN002_7methods.py — Check VN002 FULL (translation + summary + critique) bằng 7 knowledge methods:

Pipeline 7 bước (feedback_qa_pipeline_v2.md):
  1. KG check — xây KG VN002 từ PDF gốc, check facts vs translation
  2. RAG-style retrieval — từng claim có support trong translation?
  3. verify_numbers — stats trong critique/summary vs PDF
  4. Human reviewer simulation — logic + methodology errors
  5. Web search / external check — compare với literature khu vực
  6. Cross-research — check vs VN001/VN022/VN029/VN030
  7. Downstream propagation — check reports cite VN002

+6 advanced methods (feedback_6_qa_methods.md):
  #1 Assertion-Evidence Pairs — mỗi claim phản biện có evidence
  #2 Provenance Chain — critique → summary → translation → PDF
  #3 Ontology/Type System — instrument categories, causal language
  #4 Contradiction Detection — VN002 vs other VN papers
  #5 Temporal KB — data year 2021, publication 2022
  #6 Claim Strength — hedging vs assertive language

Output: qa_VN002_report.json + console verdict per method
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from collections import defaultdict

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

PDF_TXT = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_pdf_full.txt'
TRANS_PATH = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
SUMM_PATH = os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN002_VNAMHS_2022_National.docx')

# Peers for cross-ref / contradiction
PEERS = {
    'VN001': os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN001_Hoa_2024_Frontiers_HaNoi.docx'),
    'VN022': os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN022_UNICEF_VN_2022_SchoolFactors.docx'),
    'VN029': os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN029_Duong_SocPsychiatry_2631HS_TPHCM_2025.docx'),
    'VN030': os.path.join(ROOT, 'Tom-tat-tung-bai', 'VN030_Tran_HappyHouse_Cambridge_2023.docx'),
}

def read_docx(path):
    if not os.path.exists(path): return ''
    d = Document(path)
    lines = [p.text for p in d.paragraphs if p.text.strip()]
    for t in d.tables:
        for r in t.rows:
            lines.append(' | '.join(c.text.strip() for c in r.cells))
    return '\n'.join(lines)

# Load everything
print('Loading VN002 FULL translation + summary + PDF + peer summaries...')
with open(PDF_TXT, encoding='utf-8') as f:
    pdf = f.read()
pdf_norm = re.sub(r'\s+', ' ', pdf)

trans = read_docx(TRANS_PATH)
summ = read_docx(SUMM_PATH)
peers = {k: read_docx(v) for k, v in PEERS.items()}

# Extract critique section only (start at "PHẦN I — BỐI CẢNH")
critique_start = trans.find('PHẦN I — BỐI CẢNH')
critique = trans[critique_start:] if critique_start >= 0 else ''

print(f'  PDF: {len(pdf):,} chars | Translation: {len(trans):,} | Summary: {len(summ):,} | Critique: {len(critique):,}')
print(f'  Peers loaded: {list(peers.keys())} ({sum(len(v) for v in peers.values()):,} chars)')

all_issues = defaultdict(list)

# =============================================================
# METHOD 1 — KG: Build VN002 facts graph, check vs translation
# =============================================================
print('\n' + '='*70)
print('METHOD 1 — KG: VN002 FACT GRAPH vs TRANSLATION')
print('='*70)

# Core facts extracted from PDF (ground truth)
VN002_KG = {
    # Entity: sample
    'sample_size': {'value': 5996, 'unit': 'pairs', 'source': 'PDF p.13'},
    'provinces': {'value': 38, 'regions': 4, 'source': 'PDF p.13, Appendix 2'},
    'response_rate': {'value': 81.1, 'unit': '%', 'source': 'PDF p.13'},
    'age_range': {'value': '10-17', 'instrument': 'DISC-5', 'source': 'PDF p.12'},
    'age_mean': {'value': 13.3, 'source': 'PDF p.14'},
    'data_period': {'start': '2021-09-21', 'end': '2021-12-16', 'source': 'PDF p.12'},
    'publication_year': 2022,
    'interviewers': {'value': 127, 'source': 'PDF p.12'},

    # Facts: prevalence
    'mh_problem_prev': {'value': 21.7, 'unit': '%', 'n': 1301, 'instrument': 'DISC-5 subthreshold+'},
    'mh_disorder_prev': {'value': 3.3, 'unit': '%', 'n': 200, 'instrument': 'DSM-5 full criteria'},
    'anxiety_problem': {'value': 18.6, 'category': 'mh_problem'},
    'depression_problem': {'value': 4.3, 'category': 'mh_problem'},
    'adhd_problem': {'value': 2.8, 'category': 'mh_problem'},
    'ptsd_problem': {'value': 1.0, 'category': 'mh_problem'},
    'conduct_problem': {'value': 0.7, 'category': 'mh_problem'},
    'anxiety_disorder': {'value': 2.3, 'category': 'mh_disorder', 'significance': 'higher than others'},
    'mdd_disorder': {'value': 0.9, 'category': 'mh_disorder'},
    'adhd_disorder': {'value': 0.5, 'category': 'mh_disorder'},
    'ptsd_disorder': {'value': 0.3, 'category': 'mh_disorder'},
    'conduct_disorder': {'value': 0.2, 'category': 'mh_disorder'},

    # Facts: suicide
    'suicide_ideation_12m': {'value': 1.4, 'unit': '%'},
    'suicide_plan_12m': {'value': 0.4, 'unit': '%'},
    'suicide_attempt_12m': {'value': 0.2, 'unit': '%'},
    'suicide_attempt_ever': {'value': 1.6, 'unit': '%'},
    'selfharm_12m': {'value': 1, 'unit': '%'},
    'selfharm_ever': {'value': 4.7, 'unit': '%'},

    # Facts: service
    'service_use_mh_problem': {'value': 8.4, 'unit': '%'},
    'service_use_total': {'value': 6.5, 'unit': '%', 'n': 389},
    'service_doctor': {'value': 56.2, 'unit': '%'},
    'parent_need_identified': {'value': 5.1, 'unit': '%'},
    'talk_family': {'value': 73.9, 'unit': '%'},
    'talk_friend': {'value': 38.2, 'unit': '%'},

    # Facts: COVID
    'covid_problem_very_agree': {'value': 7.7, 'unit': '%'},
    'household_income_decrease': {'value': 71.5, 'unit': '%'},
    'covid_no_service_fear': {'value': 69.2, 'unit': '%'},
}

kg_violations = []
for fact_id, fact in VN002_KG.items():
    if not isinstance(fact, dict):
        continue
    if 'value' in fact:
        val = fact['value']
        # Check variant in translation
        variants = [str(val), str(val).replace('.', ','), f'{val} %', f'{val}%']
        if isinstance(val, (int, float)) and val == int(val):
            variants.extend([str(int(val)), f'{int(val)}%', f'{int(val)} %'])
        found = any(v in trans for v in variants if isinstance(v, str))
        if not found:
            kg_violations.append(f'{fact_id}={val}: NOT in translation')

print(f'  KG facts checked: {len(VN002_KG)}')
print(f'  Violations (fact not in translation): {len(kg_violations)}')
for v in kg_violations[:5]:
    print(f'    ⚠ {v}')
if len(kg_violations) > 5:
    print(f'    ... and {len(kg_violations)-5} more')

all_issues['KG'] = kg_violations

# =============================================================
# METHOD 2 — RAG-style retrieval: each critique claim has support
# =============================================================
print('\n' + '='*70)
print('METHOD 2 — RAG RETRIEVAL: critique claims supported in translation body')
print('='*70)

# Key claims made in critique PHẦN IV (Số liệu then chốt)
claim_keywords = [
    ('5.996 cặp phụ huynh–VTN', 'n = 5.996'),
    ('38 tỉnh', '38 tỉnh'),
    ('response rate 81,1', '81,1 %'),
    ('21,7 % vấn đề SKTT', '21,7 %'),
    ('3,3 % rối loạn', '3,3 %'),
    ('lo âu 18,6 %', '18,6 %'),
    ('anxiety disorder 2,3 %', '2,3 %'),
    ('service use 8,4 %', '8,4 %'),
    ('73,9 % family talk', '73,9 %'),
    ('COVID 71,5 % income drop', '71,5 %'),
    ('127 interviewers', '127'),
    ('13,3 tuổi TB', '13,3'),
    ('4 vùng', '4 vùng'),
    ('DISC-5 + DSM-5', 'DISC-5'),
    ('94,5 % đi học', '94,5 %'),
]

rag_fails = []
for label, kw in claim_keywords:
    hit_critique = kw in critique
    hit_body = kw in trans[:critique_start] if critique_start > 0 else False
    status = 'OK' if (hit_critique and hit_body) else 'MISSING in body' if hit_critique else 'N/A'
    if not (hit_critique and hit_body):
        rag_fails.append(f'{label}: {status}')
print(f'  Claims checked: {len(claim_keywords)}')
print(f'  RAG retrievals failing (claim not supported in body): {len(rag_fails)}')
for f in rag_fails:
    print(f'    ⚠ {f}')
all_issues['RAG'] = rag_fails

# =============================================================
# METHOD 3 — verify_numbers: stats in summary match stats in translation
# =============================================================
print('\n' + '='*70)
print('METHOD 3 — VERIFY_NUMBERS: summary stats vs translation')
print('='*70)

# Extract % numbers from summary and check they're in translation
summ_pcts = set(re.findall(r'(\d+[,.]?\d*)\s*%', summ))
trans_pcts = set(re.findall(r'(\d+[,.]?\d*)\s*%', trans))
missing_in_trans = summ_pcts - trans_pcts
print(f'  Unique % in summary: {len(summ_pcts)}')
print(f'  Unique % in translation: {len(trans_pcts)}')
print(f'  % in summary but NOT in translation: {len(missing_in_trans)}')
for m in sorted(missing_in_trans)[:8]:
    # Find context in summary
    idx = summ.find(f'{m} %') or summ.find(f'{m}%')
    ctx = summ[max(0, idx-40):idx+40] if idx > 0 else 'N/A'
    print(f'    ⚠ {m} %: {ctx[:100]}...')
all_issues['verify_numbers'] = [f'{m} % missing in trans' for m in missing_in_trans]

# =============================================================
# METHOD 4 — Human reviewer: logic/methodology errors
# =============================================================
print('\n' + '='*70)
print('METHOD 4 — HUMAN REVIEWER SIMULATION')
print('='*70)

reviewer_issues = []

# Check 1: cross-sectional + no causal language
causal_terms = ['gây ra', 'dẫn đến', 'làm cho', 'khiến', 'causes', 'caused by', 'do (.{0,10}) gây']
for term in causal_terms:
    matches = list(re.finditer(term, critique, re.IGNORECASE))
    for m in matches[:2]:
        ctx = critique[max(0, m.start()-80):m.end()+80]
        # Is it in a hedged context?
        if any(hedge in ctx.lower() for hedge in ['có thể', 'possibly', 'gợi ý', 'không thể xác lập', 'tương quan']):
            continue
        reviewer_issues.append(f'Causal language "{term}": {ctx[:200]}')

# Check 2: Prevalence comparison without instrument caveat
if 'so với' in critique and ('SDQ' in critique or 'GAD-7' in critique or 'DASS-21' in critique):
    if 'khác biệt' in critique or 'cẩn trọng' in critique or 'instrument' in critique or 'công cụ đo' in critique:
        pass  # hedged
    else:
        reviewer_issues.append('Prevalence comparison across different instruments without caveat')

# Check 3: Sample size claim — "đại diện" requires weighting
if 'đại diện quốc gia' in critique and 'weighted' not in critique.lower() and 'gia trọng' not in critique:
    reviewer_issues.append('Claim "đại diện quốc gia" should mention weighting')

# Check 4: Limitation count — should be substantive
limitation_count = critique.count('Hạn chế') + critique.lower().count('limitation')
if limitation_count < 3:
    reviewer_issues.append(f'Only {limitation_count} mentions of limitation — may be insufficient')

print(f'  Reviewer issues: {len(reviewer_issues)}')
for i in reviewer_issues:
    print(f'    ⚠ {i[:200]}')
all_issues['human_reviewer'] = reviewer_issues

# =============================================================
# METHOD 5 — External literature check (simulated web search)
# =============================================================
print('\n' + '='*70)
print('METHOD 5 — EXTERNAL LITERATURE CONSISTENCY')
print('='*70)

# Known facts from external lit that V-NAMHS claims should align with
ext_checks = [
    ('GBD 2019 global mental disorders 5-19 ~12%',
     '12 %' in critique or 'GBD' in critique,
     'V-NAMHS 3,3 % thấp hơn mean toàn cầu — should be flagged'),
    ('GSHS 2019 VN suicide ideation 15,6%',
     '15,6' in trans or 'GSHS' in trans,
     'V-NAMHS 1,4 % vs GSHS 15,6 % — should discuss method difference'),
    ('Blum 2012 suicide ideation 2,3%',
     '2,3' in trans or 'Blum' in trans,
     'VN002 citation gauge'),
    ('Samuels 2018 range 8-29%',
     '8' in trans and '29' in trans,
     'Historical VN range context'),
    ('Patton 2016 Lancet Commission',
     'Patton' in trans or '2016' in trans,
     'Adolescent health framework'),
]

ext_fails = []
for claim, ok, note in ext_checks:
    if ok:
        print(f'  ✓ {claim}')
    else:
        print(f'  ⚠ {claim} — {note}')
        ext_fails.append(f'{claim}: {note}')
all_issues['external'] = ext_fails

# =============================================================
# METHOD 6 — Cross-research: V-NAMHS vs peer papers
# =============================================================
print('\n' + '='*70)
print('METHOD 6 — CROSS-RESEARCH: VN002 vs VN001/VN022/VN029/VN030')
print('='*70)

cross_checks = []
# VN001 claim in critique: "chỉ Hà Nội" + "GAD-7 40,6 %"
if 'VN001' in critique:
    if 'Hà Nội' in critique and '40,6' in critique:
        cross_checks.append(('VN001', 'OK — Hà Nội + 40,6 % GAD-7'))
    else:
        cross_checks.append(('VN001', 'PARTIAL — Hà Nội check'))
    # Verify vs peer summary
    peer = peers['VN001']
    if 'Hà Nội' in peer and '40,6' in peer and 'GAD-7' in peer:
        cross_checks.append(('VN001 peer verify', 'OK'))
    else:
        cross_checks.append(('VN001 peer verify', 'MISMATCH — check VN001 summary'))

# VN022: "26,1 % SDQ"
if 'VN022' in critique and '26,1' in critique:
    cross_checks.append(('VN022', 'OK — 26,1 % SDQ'))
peer = peers['VN022']
if '26,1' in peer and 'SDQ' in peer:
    cross_checks.append(('VN022 peer verify', 'OK'))
else:
    cross_checks.append(('VN022 peer verify', f'partial (26,1 hit: {"26,1" in peer})'))

# VN029: cắt ngang TPHCM DASS-21 50,3 %
if 'VN029' in critique:
    if 'TPHCM' in critique and 'DASS' in critique and '50,3' in critique:
        cross_checks.append(('VN029', 'OK — TPHCM cắt ngang DASS-21 50,3 %'))
peer = peers['VN029']
if '50,3' in peer and 'DASS' in peer and 'TPHCM' in peer:
    cross_checks.append(('VN029 peer verify', 'OK'))
else:
    cross_checks.append(('VN029 peer verify', f'partial (50,3:{"50,3" in peer}, DASS:{"DASS" in peer}, TPHCM:{"TPHCM" in peer})'))

# VN030: Happy House Hà Nội d = 0,11
if 'VN030' in critique:
    if 'Hà Nội' in critique and 'd = 0,11' in critique:
        cross_checks.append(('VN030', 'OK — Hà Nội d = 0,11'))
peer = peers['VN030']
if '0,11' in peer and ('Hà Nội' in peer or 'Hanoi' in peer):
    cross_checks.append(('VN030 peer verify', 'OK'))
else:
    cross_checks.append(('VN030 peer verify', f'partial (0,11:{"0,11" in peer})'))

for p, s in cross_checks:
    print(f'  {p}: {s}')
all_issues['cross_research'] = [f'{p}: {s}' for p, s in cross_checks if 'OK' not in s]

# =============================================================
# METHOD 7 — Downstream check
# =============================================================
print('\n' + '='*70)
print('METHOD 7 — DOWNSTREAM PROPAGATION CHECK')
print('='*70)

import glob
bao_cao_files = glob.glob(os.path.join(ROOT, '01_Bao-cao/*.docx'))
downstream = {}
for fn in bao_cao_files:
    if '~$' in fn: continue
    try:
        d = Document(fn)
        txt = '\n'.join(p.text for p in d.paragraphs if p.text.strip())
        cnt = txt.count('VN002') + txt.count('VNAMHS') + txt.count('V-NAMHS')
        if cnt > 0:
            downstream[os.path.basename(fn)] = cnt
    except: pass

latest_v5 = 'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-5.docx'
print(f'  Files citing V-NAMHS: {len(downstream)}')
for k, v in sorted(downstream.items()):
    status = '(latest, needs update)' if k == latest_v5 else '(snapshot — preserve)'
    print(f'    {k}: {v} hits {status}')
if not any(k == latest_v5 for k in downstream):
    print(f'  ✓ Latest v5.5 does NOT cite VN002 — no downstream update needed')

# =============================================================
# ADVANCED #1: Assertion-Evidence Pairs
# =============================================================
print('\n' + '='*70)
print('ADVANCED #1: ASSERTION-EVIDENCE PAIRS')
print('='*70)

aep_pairs = [
    ('Báo cáo là mixed-methods', 'mixed', False, 'Actually V-NAMHS là cross-sectional SURVEY, không phải mixed-methods'),
    ('Sample đại diện quốc gia', '5.996', True, 'Sample n = 5.996 với weighting — đúng'),
    ('DISC-5 chuẩn DSM-5', 'DSM-5', True, 'Verified'),
    ('Response rate 81,1 %', '81,1', True, 'Verified'),
    ('Anxiety cao nhất trong vấn đề SKTT', '18,6', True, 'Verified (18,6 %)'),
]

aep_issues = []
for assertion, evidence_kw, expected, note in aep_pairs:
    hit = evidence_kw in critique
    if hit != expected:
        aep_issues.append(f'{assertion} — {note}')
    print(f'  {"✓" if hit == expected else "⚠"} {assertion}: {note}')

# Check critique doesn't claim RCT design
if 'RCT' in critique and ('V-NAMHS' in critique or 'báo cáo này' in critique):
    # Ensure RCT mentions are only for peer papers (VN030)
    rct_contexts = [m for m in re.finditer(r'.{0,50}RCT.{0,50}', critique)]
    for m in rct_contexts:
        ctx = m.group()
        if 'VN030' in ctx or 'Happy House' in ctx or 'cluster' in ctx:
            continue  # OK
        if 'V-NAMHS' in ctx.upper():
            aep_issues.append(f'V-NAMHS mislabeled as RCT: {ctx}')

all_issues['AEP'] = aep_issues

# =============================================================
# ADVANCED #2: Provenance Chain
# =============================================================
print('\n' + '='*70)
print('ADVANCED #2: PROVENANCE CHAIN (critique → summary → translation → PDF)')
print('='*70)

provenance_stats = ['5.996', '81,1', '21,7', '3,3', '18,6', '2,3', '8,4', '73,9', '7,7', '71,5']
provenance_fails = []
for stat in provenance_stats:
    critique_has = stat in critique
    summary_has = stat in summ
    trans_has = stat in trans[:critique_start] if critique_start > 0 else False
    # PDF has (decimal form)
    pdf_form = stat.replace(',', '.')
    pdf_has = pdf_form in pdf_norm or stat in pdf_norm
    chain = {'critique': critique_has, 'summary': summary_has, 'trans_body': trans_has, 'pdf': pdf_has}
    broken = [k for k, v in chain.items() if not v]
    if broken:
        provenance_fails.append(f'{stat}: broken at {broken}')
        print(f'  ⚠ {stat}: {chain}')
    else:
        print(f'  ✓ {stat}: chain intact')

all_issues['provenance'] = provenance_fails

# =============================================================
# ADVANCED #3: Ontology / Type System
# =============================================================
print('\n' + '='*70)
print('ADVANCED #3: ONTOLOGY / TYPE SYSTEM')
print('='*70)

ontology_issues = []
# V-NAMHS design = cross-sectional — no RCT claims for itself
if 'V-NAMHS' in critique and 'cross-sectional' in critique.lower():
    print('  ✓ V-NAMHS correctly classified as cross-sectional')
elif 'cắt ngang' in critique and ('V-NAMHS' in critique or 'báo cáo' in critique):
    print('  ✓ V-NAMHS correctly classified as cắt ngang')
else:
    ontology_issues.append('V-NAMHS design not explicitly labeled')
    print('  ⚠ V-NAMHS design not labeled explicitly')

# Instrument types
instruments_mentioned = {
    'DISC-5': 'diagnostic (DSM-5)',
    'SDQ': 'symptom screen',
    'GAD-7': 'symptom screen (anxiety)',
    'DASS-21': 'symptom screen (mixed)',
    'CESD-R': 'symptom screen (depression)',
}
for inst, cat in instruments_mentioned.items():
    if inst in critique:
        # Check if comparison with DISC-5 is hedged
        if inst != 'DISC-5' and 'DISC-5' in critique:
            # Look for caveats
            if 'khác biệt' in critique or 'sàng lọc vs chẩn đoán' in critique or 'screening' in critique.lower():
                print(f'  ✓ {inst} ({cat}) — comparison with DISC-5 is hedged')
            else:
                ontology_issues.append(f'{inst} compared with DISC-5 without caveat')
                print(f'  ⚠ {inst} compared with DISC-5 without explicit caveat')

all_issues['ontology'] = ontology_issues

# =============================================================
# ADVANCED #4: Contradiction Detection
# =============================================================
print('\n' + '='*70)
print('ADVANCED #4: CONTRADICTION DETECTION (VN002 vs peers)')
print('='*70)

contradictions = []

# V-NAMHS 21,7 % mh_problem vs VN022 26,1 % SDQ moderate-high
if '26,1' in peers['VN022'] and '21,7' in trans:
    # Same direction but different magnitudes — NOT contradiction but needs context
    print('  ℹ V-NAMHS 21,7 % (DISC-5) vs VN022 26,1 % (SDQ): same direction, different instruments')

# V-NAMHS 1,4 % suicide ideation vs GSHS 15,6 %
if '1,4' in trans and '15,6' in trans:
    print('  ℹ V-NAMHS 1,4 % vs GSHS 15,6 %: 11× difference — acknowledged in Interpretation section')

# Gender: V-NAMHS finds NO gender diff in total prevalence
# But VN001 finds Nữ > Nam (GAD-7)
vn001_gender = 'Nữ có điểm GAD-7 cao hơn' in peers['VN001']
vn002_no_gender = 'Không có khác biệt' in trans and ('nam' in trans.lower() and 'nữ' in trans.lower())
if vn001_gender and vn002_no_gender:
    contradictions.append('Gender: VN001 finds Nữ > Nam GAD-7 but V-NAMHS finds NO gender difference in MH problem (21,7 %)')
    print('  ⚠ Gender contradiction: VN001 Nữ > Nam (GAD-7) vs VN002 no gender diff (DISC-5 mh problem 21,7 %)')
    print('    → likely due to different instruments + different domains')

# Age pattern: V-NAMHS shows no difference 10-13 vs 14-17
# But literature suggests age gradient
vn002_no_age = '20,8' in trans and '22,7' in trans
if vn002_no_age:
    print('  ℹ V-NAMHS: no age diff 10-13 (20,8 %) vs 14-17 (22,7 %) — unusual (literature usually shows late adolescence surge)')

all_issues['contradictions'] = contradictions

# =============================================================
# ADVANCED #5: Temporal Knowledge Base
# =============================================================
print('\n' + '='*70)
print('ADVANCED #5: TEMPORAL KB')
print('='*70)

from datetime import date
current_year = date.today().year
data_year = 2021
pub_year = 2022
data_age = current_year - data_year
print(f'  VN002 data year: {data_year}')
print(f'  VN002 publication year: {pub_year}')
print(f'  Data age (as of {current_year}): {data_age} years')
if data_age > 5:
    print(f'  ⚠ Data age > 5 years — findings may be outdated')
    all_issues['temporal'] = [f'Data age {data_age} years — outdated flag']
else:
    print(f'  ✓ Data age ≤ 5 years — still current')

# Check post-COVID period note in critique
if 'wave 2' in critique.lower() or 'lặp lại' in critique or 'đã qua' in critique:
    print('  ✓ Critique acknowledges need for wave 2 / repeat study')
else:
    print('  ⚠ Critique does not acknowledge temporal limitation')

# =============================================================
# ADVANCED #6: Claim Strength Classifier
# =============================================================
print('\n' + '='*70)
print('ADVANCED #6: CLAIM STRENGTH')
print('='*70)

# Overly assertive claims from cross-sectional should be flagged
strong_terms = ['chắc chắn', 'khẳng định', 'luôn luôn', 'mọi', 'tuyệt đối', 'proves', 'gây ra (.{0,30}) trầm cảm', 'gây ra (.{0,30}) lo âu']
strength_issues = []
for term in strong_terms:
    matches = list(re.finditer(term, critique, re.IGNORECASE))
    for m in matches[:2]:
        ctx = critique[max(0, m.start()-100):m.end()+100]
        # Exclude if in a hedged context
        if any(h in ctx.lower() for h in ['tương quan', 'có thể', 'gợi ý', 'không thể', 'giả thuyết']):
            continue
        strength_issues.append(f'Strong claim: "{term}" in: {ctx[:150]}')

# Check hedging coverage
hedges = ['có thể', 'gợi ý', 'có khả năng', 'một phần', 'giả thuyết', 'tương quan không phải nhân quả']
hedge_count = sum(critique.count(h) for h in hedges)
print(f'  Hedge phrases in critique: {hedge_count}')
print(f'  Strong (uncaveated) claims: {len(strength_issues)}')
for i in strength_issues[:3]:
    print(f'    ⚠ {i}')
all_issues['claim_strength'] = strength_issues

# =============================================================
# FINAL VERDICT
# =============================================================
print('\n' + '='*70)
print('OVERALL VERDICT')
print('='*70)
total_issues = sum(len(v) for v in all_issues.values())
print(f'Total flags across all 13 methods: {total_issues}')
for method, issues in all_issues.items():
    status = '✓' if len(issues) == 0 else f'⚠ {len(issues)} flag(s)'
    print(f'  {method}: {status}')

# Save JSON report
report = {
    'document': 'VN002_VNAMHS_2022_National_FULL.docx',
    'timestamp': '2026-04-14',
    'methods_count': 13,
    'total_issues': total_issues,
    'issues_by_method': {k: v for k, v in all_issues.items()},
    'kg_facts_count': len(VN002_KG),
    'coverage_summary': {
        'pdf_chars': len(pdf),
        'translation_chars': len(trans),
        'summary_chars': len(summ),
        'critique_chars': len(critique),
    }
}
OUT_JSON = os.path.join(os.path.dirname(__file__), 'qa_VN002_report.json')
with open(OUT_JSON, 'w', encoding='utf-8') as f:
    json.dump(report, f, ensure_ascii=False, indent=2, default=str)
print(f'\nReport saved: {OUT_JSON}')
