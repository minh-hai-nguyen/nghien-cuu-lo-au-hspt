"""Build doc: Doc PRISMA flow diagram - cach doc tung buoc."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

title = d.add_heading('PRISMA flow diagram: Đọc từng bước', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Hướng dẫn đọc sơ đồ dòng chảy trong mọi bài systematic review và meta-analysis')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# META
meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
mp = cell.paragraphs[0]
mp.add_run('Câu hỏi gốc của thầy: ').bold = True
mp.add_run('"Đọc PRISMA flow là đọc như thế nào?"')
mp2 = cell.add_paragraph()
mp2.add_run('PRISMA là gì: ').bold = True
mp2.add_run('Preferred Reporting Items for Systematic reviews and Meta-Analyses. Đây là CHUẨN BÁO CÁO bắt buộc cho hầu hết tạp chí Q1 khi công bố SR/MA. Phiên bản mới nhất: PRISMA 2020 (Page et al. 2021 BMJ).')
d.add_paragraph()

# =========================================================
# PHẦN 1 — PRISMA FLOW LÀ GÌ
# =========================================================
add_heading('1. PRISMA flow diagram là gì?', level=1)

add_para('PRISMA flow diagram = SƠ ĐỒ DÒNG CHẢY thể hiện quá trình "lọc" các bài nghiên cứu từ HÀNG NGHÌN bài tìm được ban đầu xuống còn số bài CUỐI CÙNG đưa vào SR/MA.', bold=True)
d.add_paragraph()

add_para('Mục đích:')
for it in [
    '① Minh bạch — người đọc biết chính xác bao nhiêu bài bị loại, lý do gì',
    '② Tái lặp (reproducibility) — người khác có thể làm lại quy trình, ra kết quả gần giống',
    '③ Đánh giá chất lượng SR — nếu flow thiếu lý do loại bỏ, đó là cờ đỏ',
    '④ Phát hiện selection bias — nếu loại quá nhiều bài mà không rõ tại sao → nghi ngờ',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# PHẦN 2 — 4 GIAI ĐOẠN
# =========================================================
add_heading('2. Bốn giai đoạn trong PRISMA 2020 flow', level=1)

add_para('Mỗi PRISMA flow được chia thành 4 giai đoạn (stage) chính, theo thứ tự từ trên xuống:', bold=False)
d.add_paragraph()

stage_tbl = d.add_table(rows=5, cols=3)
stage_tbl.style = 'Table Grid'
hdr = stage_tbl.rows[0]
for i, h in enumerate(['Giai đoạn', 'Tiếng Anh', 'Làm gì']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

stages = [
    ('① Nhận diện', 'Identification',
     'Tìm kiếm tất cả records từ databases (PubMed, PsychINFO, Web of Science...) + nguồn khác (trích dẫn thủ công, registry). Loại DUPLICATES.'),
    ('② Sàng lọc', 'Screening',
     'Đọc TITLE và ABSTRACT, loại các bài rõ ràng không liên quan (vd: nghiên cứu trên chuột, sai tuổi, sai outcome).'),
    ('③ Đánh giá tính đủ điều kiện', 'Eligibility',
     'Đọc FULL-TEXT các bài còn lại, áp tiêu chí PICOS đầy đủ. LOẠI với LÝ DO CỤ THỂ (vd: không RCT, mẫu <10, sai outcome).'),
    ('④ Bao gồm', 'Included',
     'Các bài CUỐI CÙNG đưa vào SR. Trong đó một phần (subset) có đủ dữ liệu định lượng → đưa vào MA.'),
]
for i, (vi, en, desc) in enumerate(stages):
    row = stage_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(vi)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    cell1 = row.cells[1]
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(en)
    r1.italic = True
    row.cells[2].text = desc

d.add_paragraph()

# =========================================================
# PHẦN 3 — ĐỌC TỪNG GIAI ĐOẠN
# =========================================================
add_heading('3. Cách đọc từng giai đoạn — điểm gì cần chú ý', level=1)

add_para('3.1. Giai đoạn I — IDENTIFICATION (Nhận diện)', bold=True, size=12, color=(31, 73, 125))

id_tbl = d.add_table(rows=4, cols=2)
id_tbl.style = 'Table Grid'
ids = [
    ('Records identified from databases',
     'Tổng số bài tìm được từ mỗi database. Ví dụ Walder 2025: PubMed + PsycINFO + Scopus + Cochrane + Web of Science → 2.149 records.'),
    ('Records from other sources',
     'Nguồn khác: reference list của bài khác (backward citation search), forward citation (bài cite bài), grey literature (luận án, báo cáo), author contact.'),
    ('Duplicates removed',
     'Do cùng 1 bài xuất hiện trong nhiều database → loại trùng. Walder 2025: -474 duplicates → còn 1.675.'),
    ('ĐIỂM CẦN CHÚ Ý',
     '(a) Số database đã search (ít nhất 3); (b) Ngày cập nhật search cuối cùng; (c) Tiếng nào (Anh, Đức, Trung...); (d) Có search grey literature không.'),
]
for i, (k, v) in enumerate(ids):
    row = id_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v

d.add_paragraph()

add_para('3.2. Giai đoạn II — SCREENING (Sàng lọc)', bold=True, size=12, color=(31, 73, 125))

sc_tbl = d.add_table(rows=4, cols=2)
sc_tbl.style = 'Table Grid'
scs = [
    ('Title screening',
     'Đọc TIÊU ĐỀ, loại bài RÕ RÀNG không liên quan. Không cần đọc abstract. Thường cắt được ~60–80%.'),
    ('Abstract screening',
     'Đọc TÓM TẮT các bài còn lại. Cắt tiếp ~50% nữa.'),
    ('2 reviewer độc lập',
     'CHUẨN PRISMA: mỗi bài phải được ít nhất 2 người sàng lọc ĐỘC LẬP. Bất đồng → thảo luận hoặc người thứ 3 phán quyết. Walder 2025 ghi rõ: "NW and AF independently screened all titles and abstracts".'),
    ('ĐIỂM CẦN CHÚ Ý',
     '(a) Có 2 reviewer độc lập không; (b) Có báo cáo inter-rater reliability (kappa, %agreement) không; (c) Số bài loại có hợp lý không — loại quá ít = tiêu chí quá lỏng, loại quá nhiều = tiêu chí quá chặt.'),
]
for i, (k, v) in enumerate(scs):
    row = sc_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v

d.add_paragraph()

add_para('3.3. Giai đoạn III — ELIGIBILITY (Đánh giá tính đủ điều kiện)', bold=True, size=12, color=(31, 73, 125))

eli_tbl = d.add_table(rows=4, cols=2)
eli_tbl.style = 'Table Grid'
elis = [
    ('Full-text retrieval',
     'Tải toàn văn các bài qua sàng lọc abstract. Một số không tải được → "reports not retrieved" (paywall, luận án chưa public).'),
    ('Full-text assessment',
     'Đọc CHI TIẾT, áp từng tiêu chí PICOS. Đây là lúc loại nhiều nhất — thường chỉ ~10–25% qua được.'),
    ('LÝ DO LOẠI TỪNG BÀI',
     'QUAN TRỌNG NHẤT — PRISMA 2020 bắt buộc LIỆT KÊ LÝ DO cụ thể ở giai đoạn này. Vd Walder 2025 loại 1 bài vì "insufficient data at post-assessment". Nếu flow không có lý do → VI PHẠM PRISMA.'),
    ('ĐIỂM CẦN CHÚ Ý',
     '(a) Có liệt kê lý do exclusion không; (b) Phân loại theo PICOS (P, I, C, O, S) — bao nhiêu bài loại do sai Population, Intervention, ...; (c) Tỷ lệ loại có phù hợp câu hỏi không.'),
]
for i, (k, v) in enumerate(elis):
    row = eli_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FFE699')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(191, 97, 14)
    row.cells[1].text = v

d.add_paragraph()

add_para('3.4. Giai đoạn IV — INCLUDED (Bao gồm)', bold=True, size=12, color=(31, 73, 125))

in_tbl = d.add_table(rows=3, cols=2)
in_tbl.style = 'Table Grid'
ins = [
    ('Studies included in SR',
     'Số bài cuối cùng mô tả định tính. Đây là "mẫu số" của SR.'),
    ('Studies included in MA',
     'Số bài CÓ ĐỦ dữ liệu định lượng (mean, SD, n) để tính effect size. Thường ≤ SR. Ví dụ Walder: 22 SR, 21 MA (1 bài không đủ data post-test).'),
    ('ĐIỂM CẦN CHÚ Ý',
     '(a) SR > MA bao nhiêu bài; (b) Nếu chênh lớn → đọc lý do kỹ; (c) Tổng n (cỡ mẫu gộp) đủ power cho MA không.'),
]
for i, (k, v) in enumerate(ins):
    row = in_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = v

d.add_paragraph()

# =========================================================
# PHẦN 4 — VÍ DỤ WALDER 2025
# =========================================================
add_heading('4. Ví dụ đọc PRISMA flow Walder 2025 (QT040)', level=1, color=(54, 95, 44))

add_para('Đây là ví dụ cụ thể — áp dụng khung 4 giai đoạn cho bài đã verify từ PDF:', bold=False)
d.add_paragraph()

wal_tbl = d.add_table(rows=8, cols=3)
wal_tbl.style = 'Table Grid'
hdr = wal_tbl.rows[0]
for i, h in enumerate(['Bước', 'Số bài', 'Diễn giải']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

wal_data = [
    ('Stage I — Database search', '2.149 records', 'Tổng tìm được từ các database điện tử'),
    ('Stage I — Remove duplicates', '−474', 'Bài trùng giữa các database. Còn 1.675.'),
    ('Stage II — Title screening', 'loại 1.357', 'Đọc tiêu đề, loại bài rõ ràng không liên quan. Còn 318.'),
    ('Stage II — Abstract screening', 'loại 205', 'Đọc tóm tắt, loại tiếp. Còn 113.'),
    ('Stage III — Full-text retrieval + assess', '113 → đánh giá', 'Tải toàn văn, áp PICOS. Đa số bị loại.'),
    ('Stage III — Excluded with reasons', 'loại 91 bài', 'Lý do: sai P (tuổi >25), sai I (không phải DMHI), sai S (không RCT)...'),
    ('Stage IV — Included in SR', '22 bài', 'Mô tả tổng quan định tính'),
    ('Stage IV — Included in MA', '21 bài', '1 bài loại: Vigerland et al. [93] — "insufficient data at post-assessment"'),
]
for i, row_data in enumerate(wal_data):
    row = wal_tbl.rows[i+1] if i < 7 else None
    if row is None:
        continue
    for j, v in enumerate(row_data):
        row.cells[j].text = v
        if j == 0:
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True

# Add last row
last = wal_tbl.add_row()
for j, v in enumerate(wal_data[7]):
    last.cells[j].text = v
    if j == 0:
        for p in last.cells[j].paragraphs:
            for r in p.runs:
                r.bold = True

d.add_paragraph()

# Verify arithmetic
arith_tbl = d.add_table(rows=1, cols=1)
arith_tbl.style = 'Table Grid'
arc = arith_tbl.rows[0].cells[0]
shade_cell(arc, 'E2EFDA')
arp = arc.paragraphs[0]
arr = arp.add_run('Kiểm tra số học: ')
arr.bold = True
arr.font.color.rgb = RGBColor(54, 95, 44)
arp.add_run('2.149 − 474 − 1.357 − 205 = 113 ✓ | 113 − 91 = 22 ✓ | 22 − 1 = 21 ✓. Flow Walder 2025 CỘNG ĐÚNG → đạt chuẩn PRISMA.')
d.add_paragraph()

# =========================================================
# PHẦN 5 — DẤU HIỆU CỜ ĐỎ
# =========================================================
add_heading('5. Cờ đỏ khi đọc PRISMA flow', level=1, color=(192, 80, 77))

add_para('Khi flow có các dấu hiệu dưới đây → nghi ngờ chất lượng SR:', bold=False)
d.add_paragraph()

rf_tbl = d.add_table(rows=7, cols=2)
rf_tbl.style = 'Table Grid'
hdr = rf_tbl.rows[0]
for i, h in enumerate(['Cờ đỏ', 'Ý nghĩa']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

rfs = [
    ('Không có lý do loại bỏ ở Stage III',
     'PRISMA 2020 bắt buộc LIỆT KÊ LÝ DO. Thiếu → vi phạm chuẩn.'),
    ('Số học không cộng đúng',
     'Ví dụ 1.000 records − 200 loại = 900 nhưng báo 850 còn lại → error. Có thể có bước loại ngầm không khai báo.'),
    ('Chỉ 1 reviewer',
     'Risk cao có bias chủ quan. Chuẩn là 2 reviewer + người thứ 3 giải quyết bất đồng.'),
    ('Không search ≥ 3 database',
     'Rủi ro bỏ sót. Tối thiểu: PubMed + một database chuyên ngành + Cochrane Central.'),
    ('Chênh SR vs MA quá lớn (>30%)',
     'Ví dụ SR = 50 bài nhưng MA chỉ 25 bài → nhiều bài không đủ data → chất lượng bằng chứng yếu hoặc selection bias.'),
    ('Không cập nhật search trong 2 năm qua',
     'Bằng chứng thay đổi nhanh. Search cũ hơn 2 năm trước ngày submit = bằng chứng outdated.'),
]
for i, (k, v) in enumerate(rfs):
    row = rf_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FCE4D6')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(192, 0, 0)
    row.cells[1].text = v

d.add_paragraph()

# =========================================================
# PHẦN 6 — CHECKLIST 7 CÂU HỎI
# =========================================================
add_heading('6. Checklist 7 câu hỏi khi đọc PRISMA flow', level=1)

add_para('Khi lần đầu nhìn vào 1 flow, thầy hãy chạy qua 7 câu hỏi sau:', bold=False)
d.add_paragraph()

checklist = [
    '□ 1. Tổng records identified là bao nhiêu? (Stage I)',
    '□ 2. Bao nhiêu database đã search? Ngày cập nhật search cuối?',
    '□ 3. Số duplicates loại bỏ có hợp lý không? (thường 20-40% tổng)',
    '□ 4. Có 2 reviewer độc lập sàng lọc không? Inter-rater reliability bao nhiêu?',
    '□ 5. Có LIỆT KÊ LÝ DO LOẠI ở Stage III (full-text) không?',
    '□ 6. SR bao nhiêu bài, MA bao nhiêu bài? Chênh bao nhiêu — có lý do rõ không?',
    '□ 7. Số học có cộng đúng từ đầu đến cuối không?',
]
for c in checklist:
    p = d.add_paragraph(c)
    for r in p.runs:
        r.font.size = Pt(11)
        r.bold = True

d.add_paragraph()

# =========================================================
# PHẦN 7 — TÓM GỌN
# =========================================================
add_heading('7. Tóm gọn 3 ý cốt lõi', level=1)

core_tbl = d.add_table(rows=3, cols=1)
core_tbl.style = 'Table Grid'
cores = [
    ('① 4 giai đoạn — Identification → Screening → Eligibility → Included',
     'Đọc từ trên xuống dưới. Mỗi giai đoạn có số BÀI CÒN LẠI và số BÀI LOẠI. Quan trọng: Stage III (full-text) phải có LÝ DO LOẠI cụ thể.',
     'E2EFDA'),
    ('② SR luôn ≥ MA — MA là tập con của SR',
     'Một số bài vào SR nhưng không vào MA vì thiếu dữ liệu định lượng. Chênh 1-2 bài là bình thường; chênh >30% thì nghi ngờ.',
     'FFF2CC'),
    ('③ Kiểm tra 3 thứ: số học, lý do loại, 2 reviewer',
     'Số học phải cộng đúng. Lý do loại phải liệt kê (PRISMA 2020 yêu cầu). Phải có 2 reviewer độc lập. Thiếu 1 trong 3 = cờ đỏ về chất lượng SR.',
     'DEEBF7'),
]
for i, (k, v, clr) in enumerate(cores):
    cell = core_tbl.rows[i].cells[0]
    shade_cell(cell, clr)
    p = cell.paragraphs[0]
    r1 = p.add_run(k + '\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)

d.add_paragraph()

# TLTK
add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Page MJ, McKenzie JE, Bossuyt PM, et al. (2021). The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ, 372:n71. [Chuẩn PRISMA 2020 chính thức]',
    'Page MJ, Moher D, Bossuyt PM, et al. (2021). PRISMA 2020 explanation and elaboration. BMJ, 372:n160. [Giải thích chi tiết từng mục]',
    'Liberati A, Altman DG, Tetzlaff J, et al. (2009). The PRISMA statement for reporting systematic reviews and meta-analyses. PLoS Medicine, 6(7):e1000100. [PRISMA 2009 — phiên bản cũ]',
    'Walder N, Frey A, Berger T, et al. (2025). Digital mental health interventions for SAD in young people: SR and MA. JMIR. [QT040 — ví dụ flow chart chuẩn]',
    'Li Q, et al. (2025). Effects of Different Interventions on Anxiety Disorders in Children and Adolescents: A Bayesian NMA. BMC Psychiatry. [QT029 — flow chart đơn giản hơn]',
    'http://www.prisma-statement.org — Website chính thức, có template flow diagram download miễn phí',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 20/04/2026 | Số liệu Walder 2025 (2.149, 474, 1.357, 205, 113, 22, 21) verify trực tiếp từ PDF QT040. Khung 4 giai đoạn PRISMA dựa trên Page et al. 2021 BMJ (PRISMA 2020).')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/PRISMA_flow_cach_doc_tung_buoc.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
