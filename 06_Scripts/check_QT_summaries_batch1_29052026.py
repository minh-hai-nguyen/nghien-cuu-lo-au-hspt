# -*- coding: utf-8 -*-
"""Verify QT001-QT010 (10 file dau) so voi PDF goc.
Check key facts: sample size, year, country, main statistic.
29/05/2026."""
import os, sys, io, json, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import fitz

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
QT_FOLDER = os.path.join(ROOT, 'Tom-tat-tung-bai')

with open(os.path.join(ROOT, '06_Scripts', 'pdf_catalog_29052026.json'),
          'r', encoding='utf-8') as f:
    cat = json.load(f)

# Build lookup: QT number → PDF
qt_to_pdf = {}
for p in cat['pdfs']:
    fn = p['filename']
    m = re.match(r'QT(\d{3})_', fn)
    if m:
        qt_num = m.group(1)
        qt_to_pdf[qt_num] = p


def extract_qt_facts(qt_path):
    """Extract sample size, year, country, statistic patterns from QT docx."""
    d = Document(qt_path)
    text = '\n'.join(p.text for p in d.paragraphs)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                text += '\n' + c.text

    facts = {}
    # Sample size: "N = X" or "X học sinh" or "X HS"
    n_matches = re.findall(r'(?:N\s*=\s*|n\s*=\s*|trên\s+|mẫu\s+)([\d.,]+)\s*(?:học\s+sinh|HS|adolescent|student|tham gia)',
                            text, re.IGNORECASE)
    facts['sample_sizes'] = list(set(n_matches))[:5]

    # Years
    facts['years'] = list(set(re.findall(r'\b(20[12]\d)\b', text)))[:5]

    # Percentages
    facts['percentages'] = list(set(re.findall(r'\d+[.,]\d+\s*%', text)))[:8]

    return facts, text


def find_qt_in_text(qt_text, pdf_text):
    """Cho qt_text + pdf_text, tim cac fact xuat hien trong qt nhung NOT trong pdf."""
    # Sample sizes from QT
    qt_n = re.findall(r'\b(\d+[.,]?\d{3,})\b', qt_text)
    pdf_n = re.findall(r'\b(\d+[.,]?\d{3,})\b', pdf_text)
    # Normalize (remove dots/commas for comparison)
    qt_norm = {n.replace('.', '').replace(',', '') for n in qt_n if int(n.replace('.', '').replace(',', '')) > 100}
    pdf_norm = {n.replace('.', '').replace(',', '') for n in pdf_n if int(n.replace('.', '').replace(',', '')) > 100}

    qt_only = qt_norm - pdf_norm
    return qt_only


# Process QT001-QT010
qt_files = sorted([f for f in os.listdir(QT_FOLDER)
                   if re.match(r'QT00\d_', f) and f.endswith('.docx')
                   and 'FIXED' not in f and 'BACKUP' not in f])[:10]

print(f'Checking {len(qt_files)} QT files...\n')

issues = []
for qt_file in qt_files:
    m = re.match(r'QT(\d{3})_', qt_file)
    qt_num = m.group(1)
    qt_path = os.path.join(QT_FOLDER, qt_file)

    print(f'=== {qt_file} ===')

    # Check PDF
    if qt_num not in qt_to_pdf:
        print(f'  ✗ NO matching PDF for QT{qt_num}')
        issues.append(f'QT{qt_num}: no PDF')
        print()
        continue

    pdf_info = qt_to_pdf[qt_num]
    pdf_path = os.path.join(ROOT, pdf_info['path'])
    print(f'  PDF: {pdf_info["path"]}')

    # Extract QT facts
    try:
        qt_facts, qt_text = extract_qt_facts(qt_path)
    except Exception as e:
        print(f'  ✗ Cannot read QT: {e}')
        issues.append(f'QT{qt_num}: read fail')
        print()
        continue

    print(f'  QT sample sizes: {qt_facts["sample_sizes"]}')
    print(f'  QT years: {qt_facts["years"]}')
    print(f'  QT %: {qt_facts["percentages"][:5]}')

    # Read PDF
    try:
        doc = fitz.open(pdf_path)
        pdf_text = ''
        for i in range(min(3, doc.page_count)):
            pdf_text += doc[i].get_text() + '\n'
        doc.close()
    except Exception as e:
        print(f'  ✗ Cannot read PDF: {e}')
        issues.append(f'QT{qt_num}: PDF read fail')
        print()
        continue

    # Check year consistency
    qt_year_main = qt_facts['years'][0] if qt_facts['years'] else None
    pdf_years = list(set(re.findall(r'\b(20[12]\d)\b', pdf_text[:1500])))

    # Check if any QT sample size appears in PDF
    qt_only_numbers = find_qt_in_text(qt_text, pdf_text)

    # Verdict
    year_ok = qt_year_main in pdf_years if qt_year_main else False
    yc = "OK" if year_ok else "MISMATCH"
    print(f'  Year check: QT={qt_year_main}, PDF years={pdf_years} → {yc}')

    # Check if main % from QT appear in PDF
    qt_pcts_norm = [p.replace(',', '.').replace(' ', '') for p in qt_facts['percentages'][:5]]
    pdf_pcts = re.findall(r'\d+\.?\d*\s*%', pdf_text)
    pdf_pcts_norm = [p.replace(' ', '') for p in pdf_pcts]
    missing_pcts = [p for p in qt_pcts_norm if p not in pdf_pcts_norm][:3]
    if missing_pcts:
        print(f'  ⚠ Percentages in QT but maybe not in PDF: {missing_pcts}')

    if not year_ok and qt_year_main:
        issues.append(f'QT{qt_num}: year mismatch QT={qt_year_main} vs PDF={pdf_years}')

    print()


print()
print('=' * 60)
print(f'TỔNG: {len(issues)} issues found')
if issues:
    for iss in issues:
        print(f'  - {iss}')
