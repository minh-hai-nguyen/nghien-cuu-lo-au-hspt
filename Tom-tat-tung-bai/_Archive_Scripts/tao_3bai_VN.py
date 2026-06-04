# -*- coding: utf-8 -*-
"""3 bài Việt Nam — tóm tắt CTH v5 (≤2,5 trang/bài)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def rb(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(doc, t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)

def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ============================================================
# BÀI VN-1: HOA 2024
# ============================================================
d1 = make_doc()
bl(d1, 'Tóm tắt bài VN-1: Hoa et al. 2024', bold=True)

rb(d1, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn')
bl(d1, 'Công trình \u00ab Triệu chứng lo âu và chiến lược ứng phó ở học sinh THPT tại Việt Nam sau đại dịch COVID-19 \u00bb (Anxiety Symptoms and Coping Strategies among High School Students in Vietnam after COVID-19 Pandemic), do Phạm Thị Thu Hoa, Đỗ Thị Trang, Nguyễn Thị Liên và Ngô Anh Vinh (2024), khảo sát 3.910 học sinh từ 13 trường THPT tại Hà Nội, Việt Nam. Frontiers in Public Health, 12:1232856. DOI: 10.3389/fpubh.2024.1232856.')

rb(d1, 'Phương pháp nghiên cứu')
bl(d1, 'Công trình sử dụng Thang đo Rối loạn Lo âu Tổng quát 7 mục (GAD-7 \u2014 Generalized Anxiety Disorder 7-item Scale) kết hợp phỏng vấn sâu (in-depth interviews). Nói cách khác, đây là phương pháp hỗn hợp (mixed methods) kết hợp nghiên cứu cắt ngang và phỏng vấn định tính.', bold=True)
bl(d1, 'GAD-7 (Spitzer et al., 2006) với Cronbach alpha = 0,916 — độ tin cậy rất cao. Mẫu 3.910 HS từ 13 trường THPT Hà Nội.')

rb(d1, 'Kết quả nghiên cứu định lượng')
tbl(d1,
    ['Mức độ lo âu', 'Tỷ lệ'],
    [['Tổng có lo âu (GAD-7 ≥5)', '40,6%'],
     ['Nhẹ (5\u20139)', '23,9%'],
     ['Trung bình (10\u201314)', '10,9%'],
     ['Nặng (15\u201321)', '5,8%']],
    widths=[6.0, 3.0])
d1.add_paragraph()

tbl(d1,
    ['Chỉ số', 'Nam', 'Nữ', 'P'],
    [['Điểm GAD-7 trung bình', '1,50', '1,74', '<0,01'],
     ['Hướng', '\u2014', 'Cao hơn', 'Nữ > Nam']],
    widths=[4.0, 2.5, 2.5, 2.5])
d1.add_paragraph()

bh3(d1, 'Khác biệt giới tính : Nữ có điểm GAD-7 trung bình cao hơn nam (1,74 vs 1,50, P < 0,01) — phù hợp xu hướng quốc tế.')

rb(d1, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d1, '• Tỷ lệ 40,6% là mức cao — nhưng cần lưu ý đây là sàng lọc (GAD-7 ≥5), không phải chẩn đoán. So với V-NAMHS 2022 (chẩn đoán DISC-5: chỉ 2,3%), chênh lệch 17 lần.')
bl(d1, '• Phỏng vấn sâu phát hiện 4 nguồn lo âu: áp lực học tập, quan hệ xã hội, định kiến xã hội, kỳ vọng gia đình — tương tự Jenkins 2023 (COVID, phân biệt đối xử) và Wen 2020 (áp lực học tập OR=11,6).')

rb(d1, 'Kết luận')
bl(d1, 'Lo âu 40,6% ở 3.910 HS THPT Hà Nội sau COVID-19. Nữ > nam (P<0,01). Cần can thiệp nhắm vào áp lực học tập và kỳ vọng gia đình — hai nguồn lo âu phổ biến nhất theo phỏng vấn định tính.', bold=True)

rh2(d1, 'Phản biện')
bl(d1, 'Frontiers in Public Health (Q1, PubMed). Phương pháp hỗn hợp sáng tạo, GAD-7 alpha=0,916. Tuy nhiên, chỉ Hà Nội (đô thị) — không đại diện nông thôn/DTTS. Thiết kế cắt ngang. Phỏng vấn sâu không nêu rõ số lượng và cách chọn mẫu định tính.')

rh2(d1, 'Hướng nghiên cứu tiếp theo')
bl(d1, 'Mở rộng ra nông thôn và vùng DTTS. So sánh GAD-7 với DISC-5 trên cùng mẫu để ước tính khoảng cách sàng lọc\u2013chẩn đoán. Nghiên cứu dọc theo dõi lo âu hậu COVID.')

p = d1.add_paragraph()
r = p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50 Cao. Q1, n=3.910, GAD-7 alpha=0,916, phương pháp hỗn hợp.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d1.save('VN1_Hoa_2024.docx')
sys.stderr.write('VN1_Hoa_2024.docx OK\n')

# ============================================================
# BÀI VN-2: V-NAMHS 2022
# ============================================================
d2 = make_doc()
bl(d2, 'Tóm tắt bài VN-2: V-NAMHS 2022', bold=True)

rb(d2, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn')
bl(d2, 'Công trình \u00ab Khảo sát Quốc gia về Sức khỏe Tâm thần Thanh thiếu niên Việt Nam \u00bb (Vietnam National Adolescent Mental Health Survey \u2014 V-NAMHS), do Bệnh viện Nhi Trung ương phối hợp UNICEF và ĐH Melbourne thực hiện (2022), khảo sát 5.996 thanh thiếu niên 10\u201317 tuổi tại 38 tỉnh thành trên toàn quốc Việt Nam. Đây là khảo sát SKTT thanh thiếu niên quy mô quốc gia đầu tiên của Việt Nam.')

rb(d2, 'Phương pháp nghiên cứu')
bl(d2, 'Công trình sử dụng Phỏng vấn Chẩn đoán DISC-5 (Diagnostic Interview Schedule for Children, phiên bản DSM-5) — công cụ chẩn đoán lâm sàng, không phải sàng lọc. Nói cách khác, đây là phương pháp chẩn đoán chuẩn vàng cho tỷ lệ rối loạn thực sự, với kết quả thấp hơn nhiều so với sàng lọc.', bold=True)
bl(d2, 'Lấy mẫu đại diện quốc gia: 38 tỉnh, phân tầng theo vùng địa lý. DISC-5 đánh giá theo tiêu chí DSM-5.')

rb(d2, 'Kết quả nghiên cứu định lượng')
tbl(d2,
    ['Tình trạng', 'Tỷ lệ', 'Ghi chú'],
    [['Vấn đề SKTT chung', '21,7%', 'Bất kỳ vấn đề nào'],
     ['Lo âu (triệu chứng)', '18,6%', 'Sàng lọc, không chẩn đoán'],
     ['Rối loạn lo âu (DSM-5)', '2,3%', 'Chẩn đoán DISC-5'],
     ['Rối loạn trầm cảm chủ yếu', '0,9%', 'MDD, DSM-5'],
     ['Bất kỳ rối loạn nào', '3,3%', 'Chẩn đoán DSM-5'],
     ['Tiếp cận dịch vụ SKTT', '8,4%', 'Rất thấp']],
    widths=[4.5, 2.5, 4.5])
d2.add_paragraph()

bh3(d2, 'Khoảng cách sàng lọc\u2013chẩn đoán : Lo âu sàng lọc 18,6% vs chẩn đoán DISC-5 chỉ 2,3% — chênh lệch 8 lần. So với GAD-7 của Hoa 2024 (40,6%) và DASS-21 của Ngo Anh Vinh 2024 (54,4%), sự chênh lệch lên đến 17\u201324 lần. Đây là phát hiện quan trọng nhất cho phương pháp luận.')

rb(d2, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d2, '• Chỉ 8,4% thanh thiếu niên có vấn đề SKTT tiếp cận dịch vụ — thấp hơn Philippines (~2%, Puyat 2025) nhưng vẫn rất thấp so với Hoa Kỳ (82,6%, NSCH 2020).')
bl(d2, '• Cha mẹ chỉ nhận ra 5,1% con có vấn đề SKTT — gợi ý đào tạo phụ huynh là ưu tiên can thiệp.')

rb(d2, 'Kết luận')
bl(d2, 'V-NAMHS là khảo sát quốc gia đầu tiên về SKTT thanh thiếu niên Việt Nam. Rối loạn lo âu chẩn đoán chỉ 2,3% (DSM-5) nhưng triệu chứng lo âu sàng lọc 18,6% — khoảng cách 8 lần. Chỉ 8,4% tiếp cận dịch vụ. Cần mở rộng hệ thống SKTT học đường và đào tạo phụ huynh nhận diện sớm.', bold=True)

rh2(d2, 'Phản biện')
bl(d2, 'Cỡ mẫu quốc gia (5.996, 38 tỉnh), DISC-5/DSM-5 chuẩn vàng. Tuy nhiên, DISC-5 có thể bỏ sót triệu chứng dưới ngưỡng hội chứng. Dữ liệu 2022 có thể bị ảnh hưởng hậu COVID. Không phân tích giới tính chi tiết — GAP quan trọng.')

rh2(d2, 'Hướng nghiên cứu tiếp theo')
bl(d2, 'So sánh GAD-7 và DISC-5 trên cùng mẫu. Phân tích giới tính chi tiết. Nghiên cứu dọc theo dõi sau 2\u20133 năm. Đánh giá hiệu quả can thiệp tại trường.')

p = d2.add_paragraph()
r = p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. Khảo sát quốc gia đầu tiên, DISC-5/DSM-5, n=5.996, 38 tỉnh.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d2.save('VN2_VNAMHS_2022.docx')
sys.stderr.write('VN2_VNAMHS_2022.docx OK\n')

# ============================================================
# BÀI VN-3: PHAM 2024
# ============================================================
d3 = make_doc()
bl(d3, 'Tóm tắt bài VN-3: Pham et al. 2024', bold=True)

rb(d3, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn')
bl(d3, 'Công trình nghiên cứu về mối quan hệ giữa chăm sóc xã hội và sức khỏe tâm thần ở thanh thiếu niên tại các cơ sở hỗ trợ xã hội tại Huế, Việt Nam, do Pham và cộng sự (2024), khảo sát trên 273+273 thanh thiếu niên (nhóm can thiệp và nhóm đối chứng) tại các cơ sở hỗ trợ xã hội.')

rb(d3, 'Phương pháp nghiên cứu')
bl(d3, 'Nghiên cứu đánh giá tác động của chăm sóc cảm xúc (emotional care) và chăm sóc thể chất (physical care) đến sức khỏe tâm thần thanh thiếu niên. Nói cách khác, nghiên cứu phân tách hai loại chăm sóc để xác định loại nào có tác động thực sự.', bold=True)

rb(d3, 'Kết quả nghiên cứu định lượng')
tbl(d3,
    ['Loại chăm sóc', 'Beta (lo âu)', 'P', 'Ý nghĩa'],
    [['Chăm sóc cảm xúc', '\u20130,40', '<0,001', 'Giảm lo âu mạnh — YẾU TỐ BẢO VỆ'],
     ['Chăm sóc thể chất', 'Không có ý nghĩa', '>0,05', 'KHÔNG tác động']],
    widths=[4.0, 2.5, 2.0, 4.0])
d3.add_paragraph()

bh3(d3, 'Phát hiện quan trọng nhất : Chăm sóc CẢM XÚC (beta = \u20130,40, P < 0,001) giảm lo âu mạnh, trong khi chăm sóc THỂ CHẤT không có tác động có ý nghĩa. Điều này gợi ý can thiệp nên tập trung vào hỗ trợ cảm xúc (lắng nghe, chia sẻ, đồng cảm) hơn là chỉ đáp ứng nhu cầu vật chất.')

rb(d3, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d3, '• Phát hiện chăm sóc cảm xúc (beta = \u20130,40) tương tự Qiu et al. 2022: nuôi dạy tích cực giảm 70% nguy cơ trầm cảm (OR = 0,30) và 68% lo âu (OR = 0,32). Cả hai đều nhấn mạnh yếu tố gia đình/chăm sóc cảm xúc.')
bl(d3, '• Phát hiện chăm sóc thể chất không có ý nghĩa là bất ngờ — thách thức quan niệm "cung cấp vật chất đầy đủ = sức khỏe tâm thần tốt".')

rb(d3, 'Kết luận')
bl(d3, 'Chăm sóc cảm xúc (beta = \u20130,40) là yếu tố bảo vệ mạnh cho SKTT thanh thiếu niên, trong khi chăm sóc thể chất không có tác động. Can thiệp nên tập trung đào tạo kỹ năng chăm sóc cảm xúc cho người chăm sóc tại cơ sở hỗ trợ xã hội.', bold=True)

rh2(d3, 'Phản biện')
bl(d3, 'Phân tách rõ ràng chăm sóc cảm xúc vs thể chất — ít NC nào làm. Tuy nhiên, cỡ mẫu nhỏ (273+273), chỉ tại Huế, đối tượng đặc thù (cơ sở hỗ trợ xã hội, không phải trường học phổ thông). Thiết kế cắt ngang.')

rh2(d3, 'Hướng nghiên cứu tiếp theo')
bl(d3, 'Mở rộng sang trường học phổ thông. Nghiên cứu can thiệp đào tạo chăm sóc cảm xúc cho phụ huynh/giáo viên. Đánh giá dọc tác động dài hạn.')

p = d3.add_paragraph()
r = p.add_run('Đánh giá: \u2b50\u2b50\u2b50 Trung bình-Khá. Phân tách cảm xúc/thể chất sáng tạo nhưng cỡ mẫu nhỏ và đối tượng đặc thù.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d3.save('VN3_Pham_2024.docx')
sys.stderr.write('VN3_Pham_2024.docx OK\n')
