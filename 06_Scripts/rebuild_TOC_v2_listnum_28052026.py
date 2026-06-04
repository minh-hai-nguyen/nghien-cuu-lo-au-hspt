# -*- coding: utf-8 -*-
"""Rebuild TOC table - lay ca auto-number tu Word ListFormat.ListString.
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import win32com.client, pythoncom
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

# DO NOT restore - use current FILE state


# ============================================================
# 1. Extract headings via Word COM WITH ListString
# ============================================================
print("\n[1] Extract headings WITH auto-number...")
TEMP = os.path.join(ROOT, 'Luận án TS', '_temp_rebuild2.docx')
shutil.copy(FILE, TEMP)
pythoncom.CoInitialize()
word = None
headings = []
try:
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False; word.DisplayAlerts = False
    doc = word.Documents.Open(os.path.abspath(TEMP), ReadOnly=False)
    doc.Fields.Update(); doc.Repaginate()

    for level in [1, 2, 3, 4, 5]:
        for style_name in [f'Heading {level}', f'Tiêu đề {level}']:
            try:
                rng = doc.Content
                find = rng.Find
                find.ClearFormatting(); find.Style = doc.Styles(style_name)
                find.Text = ''; find.Forward = True; find.Wrap = 0; find.Format = True
                while find.Execute():
                    para = rng.Paragraphs.First
                    raw_text = para.Range.Text.strip()
                    # Clean control chars
                    text = ''.join(ch for ch in raw_text if ord(ch) >= 32 or ch in '\t\n').strip()
                    if para.Range.Tables.Count > 0: rng.Collapse(0); continue
                    if not text or len(text) > 250: rng.Collapse(0); continue
                    # Try get list auto-number
                    list_str = ''
                    try:
                        list_str = para.Range.ListFormat.ListString
                        list_str = ''.join(ch for ch in list_str if ord(ch) >= 32 or ch in '\t').strip()
                    except: pass
                    # Combine: if text already has number, don't double-add
                    if list_str and not re.match(r'^\d', text):
                        display = f"{list_str} {text}"
                    else:
                        display = text
                    page = para.Range.Information(3)
                    pos = para.Range.Start
                    headings.append({
                        'level': level, 'text': text, 'display': display,
                        'page': page, 'pos': pos, 'list_str': list_str,
                    })
                    rng.Collapse(0)
            except: pass
    doc.Close(SaveChanges=False)
finally:
    if word:
        try: word.Quit()
        except: pass
    pythoncom.CoUninitialize()
    if os.path.exists(TEMP):
        try: os.remove(TEMP)
        except: pass

# Dedupe
seen = set(); unique = []
for h in headings:
    k = (h['display'], h['page'])
    if k not in seen:
        seen.add(k); unique.append(h)
headings = sorted(unique, key=lambda h: h['pos'])
print(f"  {len(headings)} unique headings")
print()
print("Sample first 20:")
for h in headings[:20]:
    list_indicator = f" [list={h['list_str']!r}]" if h['list_str'] else ''
    print(f"  L{h['level']} p{h['page']}: {h['display'][:80]}{list_indicator}")


# ============================================================
# 2. Filter
# ============================================================
SKIP = {'MỤC LỤC', 'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
        'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
        'LỜI CAM ĐOAN', 'LỜI CẢM ƠN'}
filtered = []
for h in headings:
    if h['level'] > 3: continue
    if h['display'].upper().strip() in SKIP: continue
    filtered.append(h)
print(f"\nFiltered (H1-H3): {len(filtered)}")


# ============================================================
# 3. Combine "CHƯƠNG X" + title
# ============================================================
toc_entries = []
i = 0
while i < len(filtered):
    h = filtered[i]
    text = h['display'].strip()
    if re.match(r'^CHƯƠNG\s+\d+$', text):
        if i + 1 < len(filtered) and filtered[i+1]['level'] == 1 and filtered[i+1]['page'] - h['page'] < 3:
            combined = f"{text}: {filtered[i+1]['display'].strip()}"
            toc_entries.append((combined, h['page'], 1))
            i += 2; continue
    toc_entries.append((text, h['page'], h['level']))
    i += 1
print(f"Composed: {len(toc_entries)} entries")


# ============================================================
# 4. Open file, rebuild TOC table
# ============================================================
d = Document(FILE)
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb; break

if toc_table is None:
    print("KHONG TIM THAY TOC TABLE"); sys.exit(1)

# Clear all rows
tbl = toc_table._element
for tr in list(tbl.findall(qn('w:tr'))):
    tbl.remove(tr)
print(f"Cleared rows")

# Add new rows
def add_row(table, name, page, level):
    name = ''.join(ch for ch in str(name) if ord(ch) >= 32 or ch == '\t').strip()
    if not name: return
    row = table.add_row()
    cell0 = row.cells[0]
    cell0.paragraphs[0].text = ''
    p0 = cell0.paragraphs[0]
    indent = '   ' * (level - 1) if level > 1 else ''
    r0 = p0.add_run(indent + name)
    r0.font.name = 'Times New Roman'
    r0.font.size = Pt(13)
    if level == 1: r0.bold = True
    cell1 = row.cells[1]
    cell1.paragraphs[0].text = ''
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(str(page))
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    if level == 1: r1.bold = True

for name, page, level in toc_entries:
    add_row(toc_table, name, page, level)

print(f"Added {len(toc_entries)} rows")
d.save(FILE)
print(f"\nSaved: {FILE}, size: {os.path.getsize(FILE)//1024} KB")

# Verify
d2 = Document(FILE)
new_toc = None
for tb in d2.tables:
    if len(tb.rows) > 0 and 'MỞ ĐẦU' in tb.rows[0].cells[0].text:
        new_toc = tb; break
if new_toc:
    print(f"\nNew TOC: {len(new_toc.rows)} rows")
    print()
    print("First 30 rows:")
    for i, row in enumerate(new_toc.rows[:30]):
        print(f"  {i:>2}: {row.cells[0].text.strip()[:75]:<75} | {row.cells[1].text.strip()}")
