# -*- coding: utf-8 -*-
"""
Q2.5 EN+VN v3 -> v4: sửa các lỗi AN TOÀN phát hiện qua audit 16/05/2026.
- Sửa 4 trích dẫn bịa (Chen, Duong->Trúc Thanh Thái, Hoàng Trung Học, UNICEF) — dữ
  liệu đã verify (giống Bài 1).
- Bie: bổ sung đủ 10 tác giả.
- Thêm 2 trích dẫn ma đang thiếu trong danh mục: Xian (2024), Cai (2025).
- Sửa lỗi chính tả tiếng Việt: "giải rải"->"tách rời", "thám trắc"->"thăm dò",
  "một một"->"một".
- Sửa in-text Duong->Trúc Thanh Thái + UNICEF->Viện Xã hội học.
CÁC LỖI CẤU TRÚC (biến school attachment chưa khai trong Phương pháp, lỗi so sánh
β với OR, Phương pháp trùng lặp, bảng tạp chí EN chưa dịch) — KHÔNG tự sửa, nêu cho
tác giả.
"""
import os
import copy
from docx import Document
from docx.text.paragraph import Paragraph

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn"
EN_IN = os.path.join(ROOT, "Q25_SEM_Pathways_EN_v3.docx")
VN_IN = os.path.join(ROOT, "Q25_SEM_Pathways_VN_v3.docx")
EN_OUT = os.path.join(ROOT, "Q25_SEM_Pathways_EN_v4.docx")
VN_OUT = os.path.join(ROOT, "Q25_SEM_Pathways_VN_v4.docx")

CHEN = ("Chen, Z., Ren, S., He, R., Liang, Y., Tan, Y., Liu, Y., Wang, F., Shao, X., "
        "Chen, S., Liao, Y., He, Y., Li, J.-G., Chen, X., & Tang, J. (2023). Prevalence "
        "and associated factors of depressive and anxiety symptoms among Chinese "
        "secondary school students. BMC Psychiatry, 23, Article 580. "
        "https://doi.org/10.1186/s12888-023-05068-1")
THAI = ("Trúc Thanh Thái, Võ Lê Hồng Tuyết, Nguyễn Thị Trang, Đinh Văn Ngôn, Mai Lê "
        "Xuân, Trần Thị Hoài Thương, Nguyễn Thị Ngọc Bích, Huỳnh Mai Khánh Hà, Nguyễn "
        "Thị Thu An, Bùi Thị Hy Hân, & Dương Minh Cường. (2026). Unmasking the burden "
        "of mental health symptoms and risk behaviors in Vietnamese adolescents: "
        "Evidence from a multicenter cross-sectional study involving 2,631 high school "
        "students. Social Psychiatry and Psychiatric Epidemiology. {ADV} "
        "https://doi.org/10.1007/s00127-025-03043-7")
HOC = ("Hoàng Trung Học, & Nguyễn Thùy Dung. (2025). Levels of stress, anxiety, and "
       "depression in adolescents during and after the COVID-19 pandemic in Vietnam: "
       "A cross-sectional study. American Journal of Psychiatric Rehabilitation, "
       "28(1), 360–367. https://doi.org/10.69980/ajpr.v28i1.105")
BIE = ("Bie, F., Yan, X., Xing, J., Wang, L., Xu, Y., Wang, G., Wang, Q., Guo, J., "
       "Qiao, J., & Rao, Z. (2024). Rising global burden of anxiety disorders among "
       "adolescents and young adults: Trends, risk factors, and the impact of "
       "socioeconomic disparities and COVID-19 from 1990 to 2021. Frontiers in "
       "Psychiatry, 15, Article 1489427. https://doi.org/10.3389/fpsyt.2024.1489427")
CAI = ("Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for "
       "resilience in children and adolescents: A systematic review and meta-analysis "
       "of randomized controlled trials. Frontiers in Psychiatry, 16, Article 1594658. "
       "https://doi.org/10.3389/fpsyt.2025.1594658")
XIAN = ("Xian, J., Zhang, Y., & Jiang, B. (2024). Psychological interventions for "
        "social anxiety disorder in children and adolescents: A systematic review and "
        "network meta-analysis. Journal of Affective Disorders, 365, 614–627. "
        "https://doi.org/10.1016/j.jad.2024.08.097")


def setp(p, t):
    if p.runs:
        p.runs[0].text = t
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(t)


def whole(doc, sub, new):
    for p in doc.paragraphs:
        if sub in p.text:
            setp(p, new)
            return p
    return None


def part(doc, sub, new):
    n = 0
    for p in doc.paragraphs:
        if sub in p.text:
            setp(p, p.text.replace(sub, new))
            n += 1
    return n


def insert_after(anchor, text):
    new_el = copy.deepcopy(anchor._element)
    for ch in list(new_el):
        new_el.remove(ch)
    anchor._element.addnext(new_el)
    newp = Paragraph(new_el, anchor._parent)
    r = newp.add_run(text)
    if anchor.runs:
        r.font.name = anchor.runs[0].font.name
        r.font.size = anchor.runs[0].font.size
    return newp


def process(in_path, out_path, lang):
    d = Document(in_path)
    log = []
    adv = "Advance online publication." if lang == "EN" else "Xuất bản trực tuyến trước."
    unicef = (("Institute of Sociology, University of Queensland, & Johns Hopkins "
               "Bloomberg School of Public Health. (2022). Viet Nam adolescent mental "
               "health survey: Report on main findings. Institute of Sociology.")
              if lang == "EN" else
              ("Viện Xã hội học, Đại học Queensland, & Đại học Johns Hopkins. (2022). "
               "Viet Nam adolescent mental health survey (V-NAMHS): Report on main "
               "findings. Viện Xã hội học."))

    log.append(("ref Chen", whole(d, "Chen, Y., Wang, H., & Liu, M. (2023)", CHEN) is not None))
    log.append(("ref Duong->Thái", whole(d, "Duong, T. T. T., Tran, M. T., & Nguyen, H. P. (2025)",
                                          THAI.replace("{ADV}", adv)) is not None))
    log.append(("ref Hoàng Trung Học", whole(d,
        "Hoàng Trung Học, Nguyễn Thanh Bình, & Trần Thị Thu Hằng. (2025)", HOC) is not None))
    p_unicef = whole(d, "UNICEF Vietnam. (2022)", unicef)
    log.append(("ref UNICEF", p_unicef is not None))
    p_bie = whole(d, "Bie, F., Yan, X., Xing, J., Wang, L., Xu, Y., & Wang, G. (2024)", BIE)
    log.append(("ref Bie 10 tác giả", p_bie is not None))

    # thêm Cai sau Bie, Xian sau UNICEF
    if p_bie is not None:
        insert_after(p_bie, CAI)
        log.append(("thêm ref Cai", True))
    if p_unicef is not None:
        insert_after(p_unicef, XIAN)
        log.append(("thêm ref Xian", True))

    # in-text
    if lang == "EN":
        log.append(("in-text Duong->Thái", part(d, "(Dương et al., 2025)",
                                                "(Trúc Thanh Thái et al., 2026)")))
        log.append(("in-text UNICEF", part(d, "(UNICEF Vietnam, 2022)",
                                           "(Institute of Sociology et al., 2022)")))
    else:
        log.append(("in-text Duong->Thái", part(d, "(Dương và cộng sự, 2025)",
                                                "(Trúc Thanh Thái và cộng sự, 2026)")))
        log.append(("in-text UNICEF", part(d, "(UNICEF Việt Nam, 2022)",
                                           "(Viện Xã hội học và cộng sự, 2022)")))
        # lỗi chính tả VN
        log.append(("VN 'giải rải'->'tách rời'", part(d, "giải rải", "tách rời")))
        log.append(("VN 'thám trắc'->'thăm dò'", part(d, "thám trắc", "thăm dò")))
        log.append(("VN 'một một'->'một'", part(d, "một một pattern", "một pattern")))

    d.save(out_path)
    print(f"=== Q2.5 {lang} ===")
    for k, v in log:
        print(f"  {k}: {v}")
    print(f"  Saved: {out_path}\n")


process(EN_IN, EN_OUT, "EN")
process(VN_IN, VN_OUT, "VN")
print("[DONE]")
