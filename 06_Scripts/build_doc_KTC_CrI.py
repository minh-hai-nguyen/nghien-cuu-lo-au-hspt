# -*- coding: utf-8 -*-
"""Generate: Giải thích KTC (CI) vs CrI cho thầy — docx format."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Giai_thich_KTC_vs_CrI_cho_thay.docx')

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(12)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False, underline=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if underline: r.underline = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def PRun(parts):
    """Add paragraph with multiple runs (for inline formatting)."""
    p = doc.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 12))
        if opts.get('bold'): r.bold = True
        if opts.get('italic'): r.italic = True
        if opts.get('underline'): r.underline = True
        if opts.get('red'): r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif opts.get('color'): r.font.color.rgb = opts['color']
    return p

def H1(t): return P(t, bold=True, size=16, color=RGBColor(0x1F, 0x3A, 0x68), align=WD_ALIGN_PARAGRAPH.CENTER)
def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2', widths=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    return t

# ============================================================
# TIÊU ĐỀ
# ============================================================
H1('GIẢI THÍCH KÝ HIỆU KTC VÀ CrI')
P('Tài liệu tham khảo nhanh — khác biệt giữa Khoảng tin cậy tần suất (CI) và Khoảng tin cậy Bayesian (CrI) trong báo cáo nghiên cứu',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P('Người soạn: Nhóm dự án Lo âu | Tháng 04/2026',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P()

# ============================================================
# PHẦN 1
# ============================================================
H2('1. KTC (Khoảng tin cậy) — thuật ngữ tiếng Việt chung')

PRun([
    ('KTC', {'bold': True}),
    (' = ', {}),
    ('K', {'bold': True, 'underline': True}), ('hoảng ', {}),
    ('T', {'bold': True, 'underline': True}), ('in ', {}),
    ('C', {'bold': True, 'underline': True}), ('ậy. Tiếng Anh có ', {}),
    ('hai', {'bold': True}), (' khái niệm khác nhau:', {}),
])

P('• CI = Confidence Interval — Khoảng tin cậy theo trường phái tần suất (frequentist)', size=12)
P('• CrI = Credible Interval — Khoảng tin cậy theo trường phái Bayesian', size=12)
P()
PRun([
    ('Vấn đề dịch thuật: ', {'bold': True, 'italic': True}),
    ('tiếng Việt chỉ có ', {}),
    ('một ', {'bold': True}),
    ('cụm từ "khoảng tin cậy" nhưng tiếng Anh có ', {}),
    ('hai', {'bold': True}),
    (' khái niệm khác nhau về nền tảng thống kê. Do đó khi ghi "KTC 95 %" trong báo cáo, cần có context (ví dụ: Bayesian MA, frequentist RCT) để biết là CI hay CrI.', {}),
])
P()

# ============================================================
# PHẦN 2
# ============================================================
H2('2. So sánh CI vs CrI')

MakeTable(
    ['Đặc điểm', 'CI (Confidence Interval)', 'CrI (Credible Interval)'],
    [
        ('Tên tiếng Việt', 'Khoảng tin cậy (tần suất)', 'Khoảng tin cậy Bayesian'),
        ('Trường phái thống kê', 'Frequentist (tần suất chủ nghĩa)', 'Bayesian (Bayes)'),
        ('Phổ biến ở đâu',
         'Đa số NC y học, dịch tễ, RCT, meta-analysis frequentist',
         'Bayesian MA, Bayesian NMA, decision analysis'),
        ('Cách diễn giải',
         '"Nếu lặp lại NC vô hạn lần, 95 % CI sẽ chứa giá trị thật" — phức tạp, khó giải thích lâm sàng',
         '"Có 95 % xác suất giá trị thật nằm trong khoảng này" — trực quan, dễ hiểu'),
        ('Giả định cần thiết',
         'Cần n đủ lớn để xấp xỉ phân phối chuẩn',
         'Không cần mẫu lớn; dùng xác suất hậu nghiệm'),
        ('Tích hợp bằng chứng trước',
         'Không (mỗi NC độc lập)',
         'Có (prior knowledge) — phù hợp thực hành dựa trên bằng chứng tích luỹ'),
    ])
P()

# ============================================================
# PHẦN 3 — DIỄN GIẢI VÍ DỤ
# ============================================================
H2('3. Ví dụ minh hoạ: MD (KTC 95 % CrI) = –0,19 [–0,22; –0,17]')

H3('Cách đọc ĐÚNG (Bayesian CrI):')
P('"Có 95 % xác suất rằng hiệu quả can thiệp thực sự (MD — Mean Difference / khác biệt trung bình) nằm trong khoảng –0,22 đến –0,17. Điểm ước lượng tốt nhất là –0,19."',
  italic=True, color=RGBColor(0x0B, 0x5E, 0x20))
P()

H3('Cách đọc SAI (nhầm với CI tần suất):')
PRun([
    ('❌ ', {'bold': True, 'red': True}),
    ('"Nếu lặp lại NC vô hạn lần, 95 % các khoảng tính được sẽ chứa giá trị thật" — đây là định nghĩa ', {'red': True}),
    ('CI tần suất', {'bold': True, 'red': True}),
    (', KHÔNG phải CrI Bayesian.', {'red': True}),
])
P()

H3('Diễn giải thực tế cho thầy:')
P('Với NC "Zhang, Liang & Kang 2026" (Bài 56 trong báo cáo 10/04, Bayesian meta-analysis 31 RCT, n = 19.865 trẻ em/VTN) nói về CBT học đường:', italic=True)
P()
PRun([
    ('Kết quả gốc: "School-based CBT produced a small reduction in anxiety symptoms, ', {}),
    ('SMD: –0.19, 95 % CrI: –0.22 to –0.17', {'bold': True}),
    ('"', {}),
])
P()
P('Dịch đầy đủ:', bold=True)
P('"CBT học đường tạo ra hiệu ứng giảm triệu chứng lo âu nhỏ. SMD chuẩn hoá = –0,19 (KTC 95 % Bayesian: –0,22 đến –0,17). Với xác suất 95 %, hiệu quả thực sự nằm giữa –0,22 và –0,17."',
  italic=True)
P()

H3('Điểm chốt lâm sàng:')
PRun([
    ('• Cả khoảng đều nằm ', {}),
    ('bên trái số 0', {'bold': True}),
    (' → có ý nghĩa thống kê (hiệu quả thực sự khác 0).', {}),
])
PRun([
    ('• Nhưng độ lớn ', {}),
    ('rất nhỏ', {'bold': True}),
    (' (|SMD| < 0,2 theo tiêu chí Cohen) → ý nghĩa lâm sàng hạn chế — không nên triển khai universal, nên cân nhắc targeted.', {}),
])
P()

# ============================================================
# PHẦN 4 — KÝ HIỆU CHUẨN
# ============================================================
H2('4. Quy ước ký hiệu trong bản dịch của dự án Lo âu')

MakeTable(
    ['Ký hiệu viết gọn', 'Ý nghĩa đầy đủ'],
    [
        ('KTC 95 %', 'Khoảng tin cậy 95 % (chung — cần xem context để biết CI hay CrI)'),
        ('KTC 95 % CI', 'Khoảng tin cậy 95 % tần suất (frequentist) — dùng trong RCT, meta-analysis thông thường'),
        ('KTC 95 % CrI', 'Khoảng tin cậy 95 % Bayesian — dùng trong Bayesian MA, Bayesian NMA'),
        ('MD [KTC ...]', 'Mean Difference (khác biệt trung bình giữa 2 nhóm) — dùng khi các NC đo cùng thang'),
        ('SMD [KTC ...]', 'Standardized Mean Difference (khác biệt chuẩn hoá) — dùng khi các NC đo các thang khác nhau (ví dụ GAD-7 vs STAI-C vs DASS-21). Cohen: |SMD|<0,2 = trivial; 0,2–0,5 = small; 0,5–0,8 = medium; >0,8 = large.'),
        ('OR [KTC ...]', 'Odds Ratio (tỷ số odds) — thường dùng trong NC dịch tễ và case-control'),
        ('RR [KTC ...]', 'Risk Ratio (tỷ số nguy cơ tương đối) — thường dùng trong cohort'),
        ('HR [KTC ...]', 'Hazard Ratio (tỷ số nguy cơ tức thời) — dùng trong survival analysis'),
        ('β [KTC ...]', 'Hệ số hồi quy (beta) — dùng trong linear regression'),
        ('d (Cohen)', 'Cohen\'s d — hệ số hiệu ứng standard cho so sánh 2 nhóm (MD/SD)'),
        ('g (Hedges)', 'Hedges\' g — tương tự Cohen\'s d nhưng hiệu chỉnh cho cỡ mẫu nhỏ'),
    ])
P()

# ============================================================
# PHẦN 5 — KHI NÀO DÙNG CrI
# ============================================================
H2('5. Khi nào dùng CrI?')

P('CrI xuất hiện khi báo cáo khai báo là Bayesian meta-analysis (MA) hoặc Bayesian network meta-analysis (NMA). Các bài trong thư viện dự án Lo âu có dùng CrI:',
  italic=True)
P()

MakeTable(
    ['Canonical ID', 'Tác giả / Năm', 'Thiết kế Bayesian', 'Ghi chú'],
    [
        ('QT049', 'Zhang, Liang & Kang 2026', 'Bayesian MA', '31 RCT, n = 19.865; SMD anxiety –0,19 [95 % CrI: –0,22; –0,17]'),
        ('QT029', 'Li et al. 2025 (BMC)', 'Bayesian NMA', '30 RCT điều trị lo âu trẻ em; CBT SUCRA 0,66'),
        ('QT039', 'Xian, Zhang & Jiang 2024 (JAD)', 'Bayesian NMA', '30 RCT SAD ở VTN; iCBT SUCRA 71,2 % — xếp hạng 1'),
    ])
P()

P('Các bài khác trong thư viện dùng CI (frequentist) thông thường — nếu không nói là Bayesian thì mặc định là CI tần suất.',
  italic=True, color=RGBColor(0x66, 0x66, 0x66))
P()

# ============================================================
# PHẦN 6 — CÁCH GHI TRONG BÁO CÁO
# ============================================================
H2('6. Gợi ý cách viết trong báo cáo gửi thầy')

H3('Chuẩn viết đề xuất:')
PRun([
    ('✅ ', {'bold': True, 'color': RGBColor(0x0B, 0x5E, 0x20)}),
    ('"SMD = –0,19 (KTC 95 % Bayesian: –0,22 đến –0,17)"', {'italic': True}),
])
PRun([
    ('✅ ', {'bold': True, 'color': RGBColor(0x0B, 0x5E, 0x20)}),
    ('"SMD = –0,19 (95 % CrI: –0,22; –0,17)" — ngắn gọn, dùng trong bảng số liệu', {'italic': True}),
])
PRun([
    ('✅ ', {'bold': True, 'color': RGBColor(0x0B, 0x5E, 0x20)}),
    ('"SMD (KTC 95 % CrI) = –0,19 [–0,22; –0,17]" — format khoa học chuẩn', {'italic': True}),
])
P()

H3('Tránh:')
PRun([
    ('❌ ', {'bold': True, 'red': True}),
    ('"SMD = –0,19 (KTC 95 %: –0,22; –0,17)" — mơ hồ, không rõ CI hay CrI', {'italic': True, 'red': True}),
])
P()

H3('Khi cần lần đầu giới thiệu thuật ngữ CrI trong báo cáo:')
P('"CrI (Credible Interval — khoảng tin cậy Bayesian) khác với CI (Confidence Interval) thông thường ở chỗ cho phép diễn giải trực tiếp: có 95 % xác suất giá trị thật nằm trong khoảng này. Trong báo cáo này, CrI được dùng cho các kết quả từ Bayesian meta-analysis (ví dụ Zhang et al. 2026, Li et al. 2025, Xian et al. 2024), trong khi CI được dùng cho các kết quả từ frequentist meta-analysis."',
  italic=True, color=RGBColor(0x66, 0x66, 0x66))
P()

# ============================================================
# PHẦN 7 — KẾT LUẬN
# ============================================================
H2('7. Kết luận')

PRun([
    ('• ', {'bold': True}),
    ('KTC là thuật ngữ chung tiếng Việt, có thể chỉ CI (tần suất) hoặc CrI (Bayesian). Luôn ghi rõ để tránh mơ hồ.', {}),
])
PRun([
    ('• ', {'bold': True}),
    ('CrI trực quan hơn CI: diễn giải trực tiếp xác suất; phù hợp cho tư vấn lâm sàng và ra quyết định.', {}),
])
PRun([
    ('• ', {'bold': True}),
    ('Ví dụ chuẩn: ', {}),
    ('"SMD = –0,19 (KTC 95 % CrI: –0,22; –0,17)"', {'italic': True, 'bold': True}),
    (' hoặc ghi rõ trong phần phương pháp rằng "các khoảng CrI được tính theo mô hình Bayesian hierarchical".', {}),
])
PRun([
    ('• ', {'bold': True}),
    ('Khi đọc báo cáo: kiểm tra trong phần ', {}),
    ('Methods', {'italic': True}),
    (' xem meta-analysis là frequentist hay Bayesian để áp đúng cách diễn giải.', {}),
])
P()

# ============================================================
# PHỤ LỤC
# ============================================================
P()
P('─' * 70, color=RGBColor(0x99, 0x99, 0x99))
P('PHỤ LỤC: Nguồn tham khảo', bold=True, size=11)
P('• Kruschke JK. "Doing Bayesian Data Analysis" (2015, Academic Press) — chương 12 về Credible Intervals',
  size=10, color=RGBColor(0x66, 0x66, 0x66))
P('• Spiegelhalter D et al. "Bayesian Approaches to Clinical Trials and Health-Care Evaluation" (Wiley)',
  size=10, color=RGBColor(0x66, 0x66, 0x66))
P('• Gelman A et al. "Bayesian Data Analysis" (3rd ed., CRC Press)',
  size=10, color=RGBColor(0x66, 0x66, 0x66))
P('• Báo cáo Can thiệp 10/04/2026 — mục các bài Bài 29 (Li 2025), Bài 43 (Xian 2024), Bài 56 (Zhang 2026)',
  size=10, color=RGBColor(0x66, 0x66, 0x66))
P()
P('Tài liệu này do Nhóm dự án Lo âu soạn để giúp làm rõ ký hiệu thống kê khi đọc các báo cáo Bayesian meta-analysis.',
  italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT):,} bytes')
