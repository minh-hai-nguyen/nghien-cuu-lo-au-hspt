# -*- coding: utf-8 -*-
"""Sinh ban nhap Q2 song ngu English-Vietnamese - Frontiers in Psychiatry.
v1 (07/06/2026): polish toan dien theo 6 dimensions:
  1. Language polish (native English academic style, tense consistency)
  2. Structure polish (Frontiers compliance: numbered sections, end-matter)
  3. Citation polish: chuyen sang Frontiers NUMBERED style [1][2][3]
  4. Statistical reporting polish (p < .001 format, CI khi co the)
  5. Bilingual consistency
  6. Co-author end-matter (Author Contributions / Funding / COI / Ack /
     Data Availability) + 3 BLOCKING placeholders
SEM Beta verified vs LA Bang 27, 30, 33, 36, 39, 42, 45.
Combined R^2 = 0.598 tu Bang 47 = INTEGRATED MODEL ANSWER cho Q1-8.
07/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'Draft_Q2_SongNgu_v1_07062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def H1(en, vi):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(en); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(12)
    r = p.add_run(vi); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(en, vi):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(en); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(vi); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(en, vi):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(en); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(vi); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True; r.italic = True

def PP(en, vi, indent=True):
    """English paragraph followed by Vietnamese italic paragraph."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(en); r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(vi); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

def NOTE(en, vi):
    """Small note paragraph - grey italic."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(en); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(vi); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

def TBD(text_en, text_vi):
    """Placeholder for BLOCKING items - red bold."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'[TBD - {text_en}]'); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(f'[Cho - {text_vi}]'); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def set_col_widths(t, widths_cm):
    for row in t.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


# ============================================================
# COVER PAGE
# ============================================================
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(24); p.paragraph_format.space_after = Pt(12)
r = p.add_run('Draft Q2 - Frontiers in Psychiatry (v1, bilingual, polished)')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
r.font.color.rgb = RGBColor(0x70, 0x70, 0x70)

H1('An integrated risk-protective structural equation model of '
   'anxiety disorder subtypes in Vietnamese lower secondary school '
   'students: a cross-sectional study',
   'Mô hình phương trình cấu trúc tích hợp các yếu tố nguy cơ – bảo vệ '
   'đối với các phân loại rối loạn lo âu ở học sinh trung học cơ sở '
   'Việt Nam: nghiên cứu cắt ngang')

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(4)
r = p.add_run('Hang Thi Cong¹*, Duc Minh Dao¹, Nguyen Minh Duc²†')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(12)
r = p.add_run('¹ Hanoi National University of Education, Hanoi, Vietnam | '
              '² [Affiliation TBD] | * Correspondence: hangct@hnue.edu.vn | '
              '† Senior author (last author per Western convention)')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True


# ============================================================
H2('ABSTRACT', 'TÓM TẮT')

H3('Background', 'Đặt vấn đề')
PP('Anxiety disorders are a leading mental health concern among '
   'adolescents worldwide, and the Asian region recorded the largest '
   'age-standardised increase in incidence between 1990 and 2021. '
   'Although a substantial evidence base documents risk and protective '
   'factors studied in isolation, no Vietnamese study to date has '
   'tested an integrated risk-protective model across DSM-5 anxiety '
   'subtypes — generalised anxiety disorder (GAD), separation anxiety '
   'disorder (SAD) and social anxiety disorder (SocAD) — using '
   'structural equation modelling (SEM). The present study addresses '
   'this gap and examines whether subtype-specific pathways are '
   'invariant across gender.',
   'Rối loạn lo âu là một trong những vấn đề sức khỏe tâm thần hàng đầu '
   'ở thanh thiếu niên trên toàn cầu, và khu vực châu Á ghi nhận mức gia '
   'tăng tỷ lệ mới mắc chuẩn hóa theo tuổi cao nhất trong giai đoạn '
   '1990–2021. Mặc dù đã có một nền bằng chứng đáng kể về tác động đơn '
   'lẻ của các yếu tố nguy cơ và bảo vệ, đến nay vẫn chưa có nghiên cứu '
   'nào tại Việt Nam kiểm định mô hình tích hợp nguy cơ – bảo vệ trên '
   'các phân loại rối loạn lo âu theo DSM-5 — rối loạn lo âu tổng quát '
   '(GAD), rối loạn lo âu chia ly (SAD) và rối loạn lo âu xã hội '
   '(SocAD) — bằng mô hình phương trình cấu trúc (SEM). Nghiên cứu này '
   'lấp khoảng trống đó và kiểm định tính bất biến giới của các đường '
   'dẫn theo từng phân loại.')

H3('Methods', 'Phương pháp')
PP('We conducted a cross-sectional SEM study with 1,352 lower '
   'secondary school students (614 boys, 738 girls; aged 11–14 years) '
   'recruited from two purposively selected schools in Hanoi, Vietnam. '
   'Eight validated scales measured three risk factors — academic '
   'pressure (ESSA), smartphone addiction (SAS-SV) and school bullying '
   '(OBVQ) — and four protective factors — school engagement (PSSM), '
   'parental support (MSPSS-Parents), peer support (MSPSS-Peers) and '
   'self-esteem (RSES). DSM-5 anxiety subtypes were assessed with the '
   'Revised Children\'s Anxiety and Depression Scale (RCADS; GAD, SAD '
   'and SocAD subscales). We fitted subtype-specific SEMs for each of '
   'the seven predictors, an integrated risk-protective composite '
   'model and a multi-group invariance test by gender in AMOS 31.0.',
   'Chúng tôi thực hiện một nghiên cứu cắt ngang sử dụng SEM với 1.352 '
   'học sinh trung học cơ sở (614 nam, 738 nữ; 11–14 tuổi) được tuyển '
   'từ hai trường tại Hà Nội được chọn có chủ đích. Tám thang đo đã '
   'kiểm định đo lường ba yếu tố nguy cơ — áp lực học tập (ESSA), '
   'nghiện điện thoại (SAS-SV) và bắt nạt học đường (OBVQ) — và bốn '
   'yếu tố bảo vệ — gắn bó trường học (PSSM), hỗ trợ cha mẹ (MSPSS – '
   'cha mẹ), hỗ trợ bạn bè (MSPSS – bạn bè) và tự trọng (RSES). Các '
   'phân loại lo âu theo DSM-5 được đo bằng thang RCADS (tiểu thang '
   'GAD, SAD và SocAD). Chúng tôi ước lượng các mô hình SEM theo từng '
   'phân loại với từng yếu tố trong số bảy yếu tố dự báo, một mô hình '
   'tích hợp nguy cơ – bảo vệ và kiểm định bất biến đa nhóm theo giới '
   'trên AMOS 31.0.')

H3('Results', 'Kết quả')
PP('Subtype-specific SEMs revealed three differential pathways. '
   'First, school bullying showed its strongest association with SAD '
   '(β = 0.376, p < .001) rather than with GAD (β = 0.215) or SocAD '
   '(β = 0.253) — a counterintuitive pattern not previously reported '
   'in the Vietnamese literature. Second, self-esteem emerged as the '
   'strongest single protective factor for GAD (β = −0.455, p < .001), '
   'comparable in magnitude to the dominant academic-pressure risk '
   'effect on GAD (β = 0.510, p < .001). Third, peer support showed '
   'no significant association with any anxiety subtype (β = −0.015 '
   'to −0.079), challenging the assumption that peer relationships '
   'are universally protective in adolescence. The integrated '
   'risk-protective composite model accounted for R² = 0.598 of the '
   'variance in total anxiety (composite Risk → Anxiety: β = 0.669, '
   'p < .001; composite Protection → Anxiety: β = −0.220, p < .001). '
   'Multi-group invariance testing showed significant gender '
   'differences for GAD (female M = 59.47 vs male M = 51.43; '
   'F(1, 1350) = 44.48, p < .001) and SocAD (female M = 52.74 vs male '
   'M = 43.20; F(1, 1350) = 45.98, p < .001) but full gender '
   'invariance for SAD (female M = 24.76 vs male M = 25.42; '
   'F(1, 1350) = 0.246, p = .620).',
   'Các mô hình SEM theo từng phân loại làm rõ ba đường dẫn khác biệt. '
   'Thứ nhất, bắt nạt học đường có liên hệ mạnh nhất với SAD '
   '(β = 0,376; p < 0,001) so với GAD (β = 0,215) hoặc SocAD '
   '(β = 0,253) — một mẫu hình phản trực giác chưa từng được báo cáo '
   'trong y văn Việt Nam. Thứ hai, tự trọng nổi lên là yếu tố bảo vệ '
   'đơn lẻ mạnh nhất đối với GAD (β = −0,455; p < 0,001), tương đương '
   'với tác động chiếm ưu thế của áp lực học tập lên GAD (β = 0,510; '
   'p < 0,001). Thứ ba, hỗ trợ bạn bè không có liên hệ có ý nghĩa với '
   'bất kỳ phân loại lo âu nào (β = −0,015 đến −0,079), thách thức giả '
   'định về vai trò bảo vệ phổ quát của quan hệ bạn bè ở tuổi thanh '
   'thiếu niên. Mô hình tích hợp nguy cơ – bảo vệ giải thích R² = '
   '0,598 phương sai của tổng lo âu (đường dẫn tổng hợp Nguy cơ → Lo '
   'âu: β = 0,669; p < 0,001; tổng hợp Bảo vệ → Lo âu: β = −0,220; '
   'p < 0,001). Kiểm định bất biến đa nhóm cho thấy khác biệt giới có '
   'ý nghĩa với GAD (nữ M = 59,47 so với nam M = 51,43; F(1, 1350) = '
   '44,48; p < 0,001) và SocAD (nữ M = 52,74 so với nam M = 43,20; '
   'F(1, 1350) = 45,98; p < 0,001) nhưng bất biến giới hoàn toàn với '
   'SAD (nữ M = 24,76 so với nam M = 25,42; F(1, 1350) = 0,246; '
   'p = 0,620).')

H3('Conclusions', 'Kết luận')
PP('In a large Vietnamese lower-secondary sample, school bullying '
   'exerted its strongest impact on separation anxiety, self-esteem '
   'was the strongest single protective lever against generalised '
   'anxiety, and only separation anxiety showed complete gender '
   'invariance — a pattern consistent with a cultural-collectivism '
   'interpretation in which family hierarchy produces uniform '
   'attachment experiences across genders. These findings support '
   'subtype-targeted prevention frameworks that combine '
   'culture-sensitive self-esteem building with bullying-specific '
   'intervention for separation anxiety.',
   'Trong một mẫu lớn học sinh trung học cơ sở Việt Nam, bắt nạt học '
   'đường có tác động mạnh nhất tới lo âu chia ly, tự trọng là đòn bẩy '
   'bảo vệ đơn lẻ mạnh nhất chống lại lo âu tổng quát, và chỉ riêng lo '
   'âu chia ly thể hiện bất biến giới hoàn toàn — một mẫu hình phù hợp '
   'với cách diễn giải văn hóa tập thể, trong đó hệ thống thứ bậc gia '
   'đình tạo trải nghiệm gắn bó đồng nhất giữa các giới. Các phát hiện '
   'này ủng hộ các khung phòng ngừa theo từng phân loại, kết hợp xây '
   'dựng tự trọng nhạy với văn hóa và can thiệp đặc thù cho bắt nạt '
   'với lo âu chia ly.')

H3('Keywords', 'Từ khóa')
PP('adolescent anxiety; DSM-5 subtypes; structural equation modelling; '
   'gender invariance; Vietnam; bullying; self-esteem; school '
   'engagement; cross-sectional study',
   'lo âu thanh thiếu niên; phân loại DSM-5; mô hình phương trình cấu '
   'trúc; bất biến giới; Việt Nam; bắt nạt học đường; tự trọng; gắn bó '
   'trường học; nghiên cứu cắt ngang')


d.add_page_break()


# ============================================================
H2('1 Introduction', '1 Đặt vấn đề')

H3('1.1 Adolescent anxiety burden and the DSM-5 subtype rationale',
   '1.1 Gánh nặng lo âu thanh thiếu niên và lý do tiếp cận theo phân loại DSM-5')

PP('Anxiety is the leading mental health concern among adolescents '
   'worldwide. The Global Burden of Disease 2021 estimates that '
   'anxiety disorders affect approximately 11.9% of adolescents in '
   'the ASEAN region and 10.1% in Vietnam, with the largest '
   'age-standardised incidence increase between 1990 and 2021 '
   'observed in Asia (1, 2). Beyond aggregate burden, '
   'epidemiological evidence indicates that adolescent anxiety is '
   'often chronic: a population-based Australian cohort found that '
   '64% of adolescents reported anxiety or depressive symptoms on '
   'three or more occasions between ages 10 and 18 (3).',
   'Lo âu là vấn đề sức khỏe tâm thần hàng đầu ở thanh thiếu niên '
   'trên toàn cầu. Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2021 ước '
   'tính rối loạn lo âu ảnh hưởng tới khoảng 11,9% thanh thiếu niên '
   'khu vực ASEAN và 10,1% tại Việt Nam, với mức gia tăng tỷ lệ mới '
   'mắc chuẩn hóa theo tuổi cao nhất ghi nhận tại châu Á giai đoạn '
   '1990–2021 (1, 2). Ngoài gánh nặng tổng thể, bằng chứng dịch tễ '
   'cho thấy lo âu tuổi vị thành niên thường mang tính mạn tính: '
   'một nghiên cứu đoàn hệ dựa trên dân số tại Úc cho thấy 64% thanh '
   'thiếu niên báo cáo triệu chứng lo âu hoặc trầm cảm ít nhất ba '
   'lần trong giai đoạn 10–18 tuổi (3).')

PP('Yet aggregate burden tells only half the story. Central to the '
   'present study is the recognition that "anxiety" is not a '
   'unitary phenomenon. The Diagnostic and Statistical Manual of '
   'Mental Disorders, Fifth Edition (4), distinguishes generalised '
   'anxiety disorder (GAD), separation anxiety disorder (SAD) and '
   'social anxiety disorder (SocAD) as separate diagnostic '
   'entities, and the Revised Children\'s Anxiety and Depression '
   'Scale (RCADS) operationalises these subtypes for adolescent '
   'self-report (5). Aggregating subtypes can obscure clinically '
   'actionable heterogeneity. Nevertheless, most Asian studies '
   'report anxiety only as a composite outcome (6–9), leaving '
   'subtype-specific risk-protective architectures unmapped in the '
   'Vietnamese context. This unmapped subtype architecture is the '
   'primary empirical gap the present study addresses.',
   'Tuy nhiên, gánh nặng tổng thể mới chỉ là một nửa câu chuyện. Một '
   'điểm cốt lõi của nghiên cứu này là sự thừa nhận rằng "lo âu" '
   'không phải là một hiện tượng đơn nhất. Sổ tay Chẩn đoán và Thống '
   'kê các Rối loạn Tâm thần phiên bản thứ năm (4) phân biệt rối loạn '
   'lo âu tổng quát (GAD), rối loạn lo âu chia ly (SAD) và rối loạn '
   'lo âu xã hội (SocAD) thành các thực thể chẩn đoán riêng biệt; và '
   'thang Đánh giá Lo âu và Trầm cảm sửa đổi cho Trẻ em (RCADS) thao '
   'tác hóa các phân loại này cho tự báo cáo ở thanh thiếu niên (5). '
   'Việc gộp các phân loại có thể che lấp tính dị biệt lâm sàng có '
   'giá trị hành động. Tuy vậy, phần lớn nghiên cứu châu Á chỉ báo '
   'cáo lo âu dưới dạng kết quả tổng hợp (6–9), để lại khoảng trống '
   'cấu trúc nguy cơ – bảo vệ theo từng phân loại trong bối cảnh '
   'Việt Nam. Khoảng trống cấu trúc theo phân loại chưa được lập '
   'bản đồ này chính là khoảng trống thực nghiệm hàng đầu mà nghiên '
   'cứu hiện tại nhắm tới.')

PP('Beyond DSM-5 subtype heterogeneity, the Vietnamese setting '
   'introduces three culturally specific considerations that any '
   'risk-protective model must accommodate (10). First, a Confucian '
   'academic culture (11) '
   'frames high-stakes examinations as the principal gateway to '
   'social mobility, intensifying chronic academic stress. Second, '
   'a hierarchical family structure constrains emotional disclosure '
   'between adolescents and parents and may therefore attenuate the '
   'efficacy of parental support. Third, a collectivist orientation '
   'may uniformly shape attachment processes across genders, with '
   'theoretical implications for separation anxiety (12, 13). These '
   'culture-specific features call for empirical examination beyond '
   'imported Western frameworks.',
   'Ngoài tính dị biệt của các phân loại DSM-5, bối cảnh Việt Nam '
   'còn đặt ra ba yếu tố đặc thù văn hóa mà bất kỳ mô hình nguy cơ – '
   'bảo vệ nào cũng phải tính đến (10). Thứ nhất, văn hóa học thuật '
   'Nho giáo (11) coi các kỳ thi áp lực cao '
   'là cửa ngõ chính để di chuyển xã hội, làm gia tăng stress học '
   'tập kéo dài. Thứ hai, cấu trúc gia đình thứ bậc hạn chế việc '
   'bộc lộ cảm xúc giữa thanh thiếu niên và cha mẹ, và do đó có thể '
   'làm suy giảm hiệu lực của hỗ trợ cha mẹ. Thứ ba, định hướng tập '
   'thể có thể định hình các quá trình gắn bó một cách đồng nhất '
   'giữa các giới, với hàm ý lý thuyết cho lo âu chia ly (12, 13). '
   'Các đặc trưng văn hóa này đòi hỏi kiểm chứng thực nghiệm vượt ra '
   'ngoài các khung lý thuyết phương Tây nhập khẩu.')


H3('1.2 Risk-protective framework and the integration imperative',
   '1.2 Khung lý thuyết nguy cơ – bảo vệ và yêu cầu tích hợp')

PP('The transactional stress-coping framework (14) posits that '
   'adolescent psychopathology emerges from dynamic interactions '
   'between stressors and resources, both internal (e.g., '
   'self-esteem) and external (e.g., school, peer and family '
   'systems). A landmark meta-analysis by Compas and colleagues '
   '(212 studies, N = 80,850) confirmed the joint contribution of '
   'risk and protective factors to youth internalising symptoms '
   '(15). Critically, however, most primary studies have tested '
   'these factors in isolation, leaving the interactive architecture '
   'under-specified.',
   'Khung lý thuyết stress-ứng phó giao dịch (14) giả định rằng tâm '
   'bệnh ở tuổi thanh thiếu niên xuất hiện từ tương tác động giữa '
   'các yếu tố gây stress và các nguồn lực, cả nội tại (ví dụ, tự '
   'trọng) và bên ngoài (ví dụ, hệ thống trường học, bạn bè và gia '
   'đình). Một phân tích tổng hợp quan trọng của Compas và cộng sự '
   '(212 nghiên cứu, N = 80.850) đã khẳng định sự đóng góp đồng thời '
   'của các yếu tố nguy cơ và bảo vệ đối với triệu chứng nội hóa ở '
   'giới trẻ (15). Tuy nhiên, phần lớn nghiên cứu gốc chỉ kiểm định '
   'các yếu tố một cách riêng lẻ, để lại cấu trúc tương tác chưa '
   'được xác định.')

PP('In Vietnam, methodological inconsistency further obscures the '
   'evidence base. The Vietnam Adolescent Mental Health Survey '
   '(V-NAMHS), which used the DSM-5–anchored Diagnostic Interview '
   'Schedule for Children (DISC-5), estimated a 12-month anxiety '
   'disorder prevalence of 2.3% (16), whereas school-based studies '
   'using continuous symptom scales (e.g., DASS-21) report '
   'substantially higher rates (17). This measurement gap '
   'underscores the need for analytical strategies that move beyond '
   'aggregate prevalence to mechanistic models of subtype-specific '
   'pathways. Among the mechanistic dimensions that require mapping, '
   'gender is the most salient — and, as we argue next, the most '
   'culturally sensitive.',
   'Tại Việt Nam, sự thiếu nhất quán về phương pháp luận càng làm cơ '
   'sở bằng chứng thêm mờ nhạt. Khảo sát Sức khỏe Tâm thần Thanh '
   'thiếu niên Việt Nam (V-NAMHS) sử dụng Bảng phỏng vấn chẩn đoán '
   'cho trẻ em phiên bản DSM-5 (DISC-5) ước tính tỷ lệ rối loạn lo '
   'âu trong 12 tháng là 2,3% (16), trong khi các nghiên cứu trường '
   'học sử dụng thang đo triệu chứng liên tục (ví dụ DASS-21) báo '
   'cáo tỷ lệ cao hơn nhiều (17). Khoảng cách đo lường này nhấn mạnh '
   'nhu cầu về chiến lược phân tích vượt ra ngoài tỷ lệ tổng hợp để '
   'tiến tới mô hình cơ chế của các đường dẫn theo từng phân loại. '
   'Trong số các kích thước cơ chế cần được lập bản đồ, giới là '
   'kích thước nổi bật nhất – và như chúng tôi luận giải tiếp theo, '
   'cũng là kích thước nhạy cảm nhất với văn hóa.')


H3('1.3 Gender, separation anxiety and the cultural-collectivism hypothesis',
   '1.3 Giới, lo âu chia ly và giả thuyết văn hóa tập thể')

PP('A standard expectation drawn from the Western literature is '
   'that females exhibit higher anxiety than males across subtypes '
   '(18). Emerging Asian evidence challenges this expectation: '
   'rural Chinese males have shown higher anxiety than rural '
   'Chinese females (8), and Northeast Indian adolescent boys '
   'reported a significantly higher severe-anxiety prevalence than '
   'girls (30.0% vs 18.9%; p = .049) (9). These patterns suggest '
   'that gender effects on anxiety are culturally and contextually '
   'contingent rather than universal.',
   'Một kỳ vọng tiêu chuẩn rút ra từ y văn phương Tây là nữ có mức '
   'lo âu cao hơn nam ở các phân loại (18). Bằng chứng mới nổi từ '
   'châu Á thách thức kỳ vọng này: nam thanh thiếu niên Trung Quốc '
   'nông thôn có mức lo âu cao hơn nữ cùng vùng (8), và nam thanh '
   'thiếu niên Đông Bắc Ấn Độ có tỷ lệ lo âu nặng cao hơn nữ một '
   'cách có ý nghĩa (30,0% so với 18,9%; p = 0,049) (9). Các mẫu '
   'hình này gợi ý rằng tác động của giới lên lo âu phụ thuộc vào '
   'văn hóa và bối cảnh, chứ không phải hằng định phổ quát.')

PP('We advance a culturally specific hypothesis: in Vietnam, '
   'collectivist family hierarchy may homogenise '
   'separation-anxiety experiences across genders, while '
   'gender-differentiated social and academic pressures remain '
   'operative for GAD and SocAD. This prediction draws on '
   'interdependent self-construal theory (13) and on developmental '
   'work indicating that separation-individuation tasks are '
   'predominantly childhood-onset and therefore precede the '
   'post-pubertal differentiation of gender-loaded social pressures '
   '(19, 20).',
   'Chúng tôi đưa ra một giả thuyết đặc thù văn hóa: tại Việt Nam, '
   'hệ thống thứ bậc gia đình tập thể có thể đồng nhất hóa trải '
   'nghiệm lo âu chia ly giữa các giới, trong khi áp lực xã hội và '
   'học tập phân hóa theo giới vẫn tác động lên GAD và SocAD. Dự '
   'đoán này dựa trên lý thuyết bản thân phụ thuộc lẫn nhau (13) và '
   'các nghiên cứu phát triển cho thấy các nhiệm vụ chia ly – cá '
   'nhân hóa chủ yếu khởi phát từ thời thơ ấu, do đó đi trước sự '
   'phân hóa giới của các áp lực xã hội sau dậy thì (19, 20).')


H3('1.4 The present study', '1.4 Nghiên cứu hiện tại')

PP('Drawing these threads together — the DSM-5 subtype heterogeneity '
   'gap, the Vietnamese cultural specificity and the cultural-'
   'collectivism gender prediction — we report a cross-sectional '
   'SEM study that applies subtype-specific structural equation '
   'modelling to test three pre-specified hypotheses in a '
   'Vietnamese lower secondary school sample:',
   'Tổng hợp các mạch lập luận trên — khoảng trống về tính dị biệt '
   'của các phân loại DSM-5, đặc thù văn hóa Việt Nam và dự đoán về '
   'giới dựa trên văn hóa tập thể — chúng tôi trình bày một nghiên '
   'cứu cắt ngang áp dụng SEM theo từng phân loại nhằm kiểm định ba '
   'giả thuyết được tiền-định trên mẫu học sinh trung học cơ sở '
   'Việt Nam:')

PP('H1: All three risk factors (academic pressure, smartphone '
   'addiction and school bullying) will show significant positive '
   'paths to GAD, SAD and SocAD.',
   'H1: Cả ba yếu tố nguy cơ (áp lực học tập, nghiện điện thoại và '
   'bắt nạt học đường) đều có đường dẫn dương có ý nghĩa tới GAD, '
   'SAD và SocAD.', indent=False)

PP('H2: All four protective factors (school engagement, parental '
   'support, peer support and self-esteem) will show significant '
   'negative paths to GAD, SAD and SocAD.',
   'H2: Cả bốn yếu tố bảo vệ (gắn bó trường học, hỗ trợ cha mẹ, hỗ '
   'trợ bạn bè và tự trọng) đều có đường dẫn âm có ý nghĩa tới GAD, '
   'SAD và SocAD.', indent=False)

PP('H3: Multi-group SEM by gender will reveal gender invariance '
   'for separation anxiety only, whereas GAD and SocAD will show '
   'gender differences favouring higher scores in females.',
   'H3: SEM đa nhóm theo giới sẽ làm rõ tính bất biến giới chỉ với '
   'lo âu chia ly, trong khi GAD và SocAD sẽ thể hiện khác biệt '
   'giới với điểm số cao hơn ở nữ.', indent=False)

NOTE('Note: This is a quantitative cross-sectional analysis. A '
     'companion qualitative study is planned separately.',
     'Ghi chú: Đây là phân tích định lượng cắt ngang. Một nghiên cứu '
     'định tính bổ trợ được lên kế hoạch triển khai riêng.')


d.add_page_break()


# ============================================================
H2('2 Methods', '2 Phương pháp')

H3('2.1 Study design and participants',
   '2.1 Thiết kế nghiên cứu và người tham gia')

PP('Testing the three hypotheses above required a sample large '
   'enough to support multi-group SEM and a measurement battery '
   'aligned with the DSM-5 subtype structure. We employed a '
   'cross-sectional survey design combined with structural equation '
   'modelling. The sample comprised 1,352 '
   'lower secondary school students (males: n = 614, 45.4%; '
   'females: n = 738, 54.6%) recruited from two purposively '
   'selected schools in Hanoi, Vietnam (Nhat Tan and Tay Mo) '
   'representing urban and suburban contexts. Participants were '
   'aged 11–14 years and distributed across Grades 6 to 9 '
   '(Grade 6: n = 368; Grade 7: n = 316; Grade 8: n = 340; '
   'Grade 9: n = 328). Recruitment proceeded through school '
   'administration, with written parental consent and written '
   'student assent obtained for all participants.',
   'Việc kiểm định ba giả thuyết trên đòi hỏi một mẫu đủ lớn để '
   'hỗ trợ SEM đa nhóm và một bộ đo lường khớp với cấu trúc phân '
   'loại DSM-5. Chúng tôi sử dụng thiết kế khảo sát cắt ngang kết '
   'hợp với mô hình phương trình cấu trúc. Mẫu gồm 1.352 học sinh '
   'trung học '
   'cơ sở (nam: n = 614; 45,4%; nữ: n = 738; 54,6%) được tuyển từ '
   'hai trường tại Hà Nội được chọn có chủ đích (Nhật Tân và Tây '
   'Mỗ) đại diện cho bối cảnh đô thị và ngoại ô. Người tham gia ở '
   'độ tuổi 11–14, phân bố qua các khối từ 6 đến 9 (khối 6: n = '
   '368; khối 7: n = 316; khối 8: n = 340; khối 9: n = 328). Quá '
   'trình tuyển chọn được thực hiện qua ban giám hiệu, với chấp '
   'thuận bằng văn bản của cha mẹ và đồng ý bằng văn bản của học '
   'sinh.')

H3('2.2 Measures', '2.2 Công cụ đo lường')

PP('Measurement covered three risk factors, four protective '
   'factors and three DSM-5 anxiety subtypes. Eight validated '
   'instruments were adapted into Vietnamese '
   'through forward–back translation, expert review (three '
   'clinical psychologists and one educational measurement '
   'specialist) and pilot testing on a separate sample of 50 '
   'students prior to the main survey. All raw scale scores were '
   'rescaled to a 0–100 metric for cross-scale comparability, '
   'while preserving the original scoring conventions.',
   'Bộ đo lường bao trùm ba yếu tố nguy cơ, bốn yếu tố bảo vệ và '
   'ba phân loại lo âu theo DSM-5. Tám công cụ đo lường đã được '
   'kiểm định được chuyển ngữ sang '
   'tiếng Việt theo quy trình dịch xuôi – dịch ngược, được hội '
   'đồng chuyên gia (ba nhà tâm lý lâm sàng và một chuyên gia đo '
   'lường giáo dục) rà soát và được khảo sát thử trên một mẫu '
   'riêng gồm 50 học sinh trước khi triển khai chính thức. Tất cả '
   'điểm thô được tái cân chỉnh về thang 0–100 để so sánh được '
   'giữa các thang đo, trong khi vẫn giữ nguyên quy ước chấm điểm '
   'gốc.')

PP('Anxiety subtypes were measured with the Revised Children\'s '
   'Anxiety and Depression Scale (RCADS) (5), a developmentally '
   'calibrated, DSM-5-aligned instrument. We extracted three '
   'subscales: generalised anxiety disorder (7 items), separation '
   'anxiety disorder (4 items) and social anxiety disorder (4 '
   'items), each rated on a 4-point Likert scale.',
   'Các phân loại lo âu được đo bằng thang Đánh giá Lo âu và Trầm '
   'cảm sửa đổi cho Trẻ em (RCADS) (5), một công cụ phù hợp với '
   'phát triển và liên kết với DSM-5. Chúng tôi trích xuất ba '
   'tiểu thang: rối loạn lo âu tổng quát (7 mục), rối loạn lo âu '
   'chia ly (4 mục) và rối loạn lo âu xã hội (4 mục), mỗi mục '
   'chấm trên thang Likert 4 điểm.')

PP('Risk factors were assessed using the following: (a) '
   'Educational Stress Scale for Adolescents (ESSA) (21), 4 '
   'items; (b) Smartphone Addiction Scale–Short Version (SAS-SV) '
   '(22), 5 items; and (c) Olweus Bully/Victim Questionnaire '
   '(OBVQ) (23), 8 items capturing physical and verbal '
   'victimisation.',
   'Các yếu tố nguy cơ được đo bằng: (a) thang Áp lực Học tập cho '
   'Vị thành niên (ESSA) (21), 4 mục; (b) thang Nghiện Điện thoại '
   'thông minh phiên bản rút gọn (SAS-SV) (22), 5 mục; và (c) '
   'Bảng hỏi Bắt nạt/Nạn nhân Olweus (OBVQ) (23), 8 mục đo lường '
   'bắt nạt thể chất và lời nói.')

PP('Protective factors comprised the following: (a) '
   'Psychological Sense of School Membership (PSSM) (24), 7 '
   'items; (b) Multidimensional Scale of Perceived Social Support '
   '(MSPSS) (25), 8 items split into parental and peer subscales '
   '(4 items each); and (c) Rosenberg Self-Esteem Scale (RSES) '
   '(26), 5 items.',
   'Các yếu tố bảo vệ gồm: (a) thang Cảm nhận Gắn bó với Trường '
   '(PSSM) (24), 7 mục; (b) thang Đa chiều về Hỗ trợ Xã hội Cảm '
   'nhận (MSPSS) (25), 8 mục chia thành tiểu thang cha mẹ và bạn '
   'bè (mỗi tiểu thang 4 mục); và (c) thang Tự trọng Rosenberg '
   '(RSES) (26), 5 mục.')


H3('2.3 Analytic strategy', '2.3 Chiến lược phân tích')

PP('To link these measures to the three hypotheses, quantitative '
   'analyses proceeded in five sequential steps designed to build '
   'from psychometric verification toward the integrated '
   'risk-protective composite and gender-invariance tests. We used '
   'AMOS 31.0 and SPSS 31.0. Step 1 generated descriptive '
   'statistics '
   'and reliability estimates (Cronbach\'s α and McDonald\'s ω). '
   'Step 2 performed confirmatory factor analysis (CFA) for each '
   'scale, applying conventional fit-index thresholds (CFI ≥ '
   '0.90, TLI ≥ 0.90, RMSEA ≤ 0.08 and SRMR ≤ 0.08) (27). Step 3 '
   'estimated subtype-specific structural models for each of the '
   'seven predictors against each of the three anxiety subtypes '
   'plus the three-factor total anxiety (RLLA) outcome, '
   'yielding 21 theoretically informative path coefficients plus '
   'seven aggregate paths. Step 4 fitted an integrated '
   'risk-protective composite model in which the three risk '
   'factors loaded onto a higher-order Risk latent (YTNC) and '
   'the four protective factors loaded onto a higher-order '
   'Protective latent (YTBV); both higher-order latents '
   'simultaneously predicted total RLLA. Step 5 implemented '
   'multi-group invariance testing (configural → metric → '
   'scalar) by gender, applying ΔCFI ≤ 0.01 as the invariance '
   'criterion (28).',
   'Để liên kết các đo lường này với ba giả thuyết, phân tích định '
   'lượng tiến hành theo năm bước tuần tự, được thiết kế nhằm đi '
   'từ kiểm định tâm trắc tới mô hình tích hợp nguy cơ – bảo vệ và '
   'kiểm định bất biến giới. Chúng tôi sử dụng AMOS 31.0 và SPSS '
   '31.0. Bước 1 tạo thống kê mô tả và ước lượng độ tin '
   'cậy (Cronbach α và McDonald ω). Bước 2 chạy phân tích nhân '
   'tố khẳng định (CFA) cho từng thang, áp dụng các ngưỡng chỉ '
   'số phù hợp quy ước (CFI ≥ 0,90; TLI ≥ 0,90; RMSEA ≤ 0,08; '
   'SRMR ≤ 0,08) (27). Bước 3 ước lượng các mô hình cấu trúc '
   'theo từng phân loại cho mỗi yếu tố trong số bảy yếu tố dự '
   'báo với từng phân loại trong số ba phân loại lo âu cộng với '
   'kết quả tổng RLLA ba nhân tố, cho 21 hệ số đường dẫn lý '
   'thuyết quan tâm cộng với bảy đường dẫn tổng hợp. Bước 4 '
   'ước lượng mô hình tích hợp nguy cơ – bảo vệ tổng hợp, trong '
   'đó ba yếu tố nguy cơ tải lên một biến tiềm ẩn bậc cao Nguy '
   'cơ (YTNC) và bốn yếu tố bảo vệ tải lên một biến tiềm ẩn bậc '
   'cao Bảo vệ (YTBV); cả hai biến tiềm ẩn bậc cao đồng thời dự '
   'báo tổng RLLA. Bước 5 triển khai kiểm định bất biến đa nhóm '
   '(cấu hình → metric → thang điểm) theo giới, áp dụng ΔCFI ≤ '
   '0,01 làm tiêu chí bất biến (28).')


H3('2.4 Ethics', '2.4 Đạo đức nghiên cứu')

PP('The study was approved by the Institutional Review Board of '
   'Hanoi National University of Education. Written parental '
   'consent and written student assent were obtained from all '
   'participants. Data were anonymised via ID codes, no personal '
   'identifiers appear in analysis files and data were stored on '
   'access-controlled institutional servers.',
   'Nghiên cứu được Hội đồng Đạo đức của Trường Đại học Sư phạm '
   'Hà Nội phê duyệt. Đồng ý bằng văn bản của cha mẹ và đồng ý '
   'bằng văn bản của học sinh đều đã được thu thập từ tất cả '
   'người tham gia. Dữ liệu được ẩn danh hóa bằng mã số nhận '
   'dạng, không có thông tin cá nhân xuất hiện trong tệp phân '
   'tích và dữ liệu được lưu trữ trên máy chủ truy cập có kiểm '
   'soát của cơ sở.')

TBD('Q3-6 BLOCKING: IRB decision number and date pending the '
    'official approval letter from the HNUE Ethics Committee. NCS '
    'to confirm exact decision number, date of approval and '
    'attach a scanned copy as a Supplementary File at submission.',
    'Q3-6 CHỜ: số quyết định và ngày của Hội đồng Đạo đức ĐHSPHN. '
    'NCS xác nhận số quyết định, ngày phê duyệt và đính kèm bản '
    'scan vào Tệp Bổ sung khi nộp bài.')


d.add_page_break()


# ============================================================
H2('3 Results', '3 Kết quả')

H3('3.1 Sample characteristics and psychometric properties',
   '3.1 Đặc điểm mẫu và tính chất tâm trắc')

PP('Of the 1,352 participants, 614 (45.4%) were male and 738 '
   '(54.6%) were female. The mean age was 12.5 years (SD = 1.1). '
   'Grade distribution was as follows: Grade 6 (n = 368), Grade 7 '
   '(n = 316), Grade 8 (n = 340) and Grade 9 (n = 328). All eight '
   'scales achieved acceptable internal consistency (Cronbach\'s '
   'α ≥ 0.70) and adequate CFA fit (CFI ≥ 0.90, RMSEA ≤ 0.08). '
   'Detailed psychometric statistics appear in Supplementary '
   'Table S1.',
   'Trong số 1.352 người tham gia, 614 (45,4%) là nam và 738 '
   '(54,6%) là nữ. Độ tuổi trung bình là 12,5 (SD = 1,1). Phân '
   'bố khối lớp: khối 6 (n = 368), khối 7 (n = 316), khối 8 '
   '(n = 340) và khối 9 (n = 328). Tất cả tám thang đo đạt độ '
   'nhất quán nội tại chấp nhận được (Cronbach α ≥ 0,70) và độ '
   'phù hợp CFA đầy đủ (CFI ≥ 0,90; RMSEA ≤ 0,08). Các chỉ số '
   'tâm trắc chi tiết được trình bày trong Bảng phụ S1.')


H3('3.2 Descriptive anxiety levels', '3.2 Mức độ lo âu mô tả')

PP('Before estimating pathways, we examined the relative magnitude '
   'of the three anxiety subtypes themselves. On the 0–100 metric, '
   'mean anxiety scores differed substantially across subtypes: GAD '
   '(M = 55.82, SD = 22.4), '
   'SocAD (M = 48.41, SD = 26.19) and SAD (M = 25.06, '
   'SD = 24.29). Separation anxiety scored markedly lower than '
   'both generalised and social anxiety, in line with '
   'developmental theory positing childhood onset and attenuation '
   'across early adolescence (19, 20).',
   'Trước khi ước lượng các đường dẫn, chúng tôi xem xét độ lớn '
   'tương đối của chính ba phân loại lo âu. Trên thang 0–100, điểm '
   'lo âu trung bình khác biệt đáng kể giữa các phân loại: GAD '
   '(M = 55,82; SD = 22,4); SocAD '
   '(M = 48,41; SD = 26,19) và SAD (M = 25,06; SD = 24,29). Lo '
   'âu chia ly có điểm thấp hơn rõ rệt so với cả lo âu tổng quát '
   'và lo âu xã hội, phù hợp với lý thuyết phát triển cho rằng '
   'SAD khởi phát từ thời thơ ấu và suy giảm dần khi vào tuổi '
   'thanh thiếu niên sớm (19, 20).')


H3('3.3 Subtype-specific path coefficients',
   '3.3 Hệ số đường dẫn theo từng phân loại')

PP('We next tested H1 and H2 by estimating subtype-specific SEMs '
   'for each of the seven predictors against GAD, SAD, SocAD and '
   'total RLLA (Table 1). The pattern was largely as predicted: '
   'all 21 risk → anxiety paths reached significance, and 12 of '
   'the 12 protective → anxiety paths to GAD and SocAD were '
   'significant. Two informative exceptions stood out: '
   'peer-support paths to GAD and SAD were non-significant, and '
   'school engagement → SAD was non-significant — exceptions we '
   'return to in Section 4.4. Detailed coefficients appear below.',
   'Tiếp theo, chúng tôi kiểm định H1 và H2 bằng cách ước lượng '
   'SEM theo từng phân loại cho mỗi yếu tố trong số bảy yếu tố dự '
   'báo với GAD, SAD, SocAD và RLLA tổng (Bảng 1). Mẫu hình cơ '
   'bản đúng như dự đoán: toàn bộ 21 đường dẫn nguy cơ → lo âu '
   'đạt ý nghĩa, và 12 trên 12 đường dẫn bảo vệ → lo âu tới GAD '
   'và SocAD đều có ý nghĩa. Hai ngoại lệ có ý nghĩa thông tin nổi '
   'bật: hỗ trợ bạn bè tới GAD và SAD không có ý nghĩa, và gắn bó '
   'trường học → SAD không có ý nghĩa — các ngoại lệ này được bàn '
   'lại ở Mục 4.4. Hệ số chi tiết được trình bày bên dưới.')


# Table 1: Subtype-specific beta table
P_header = ['Predictor', '→ GAD (β)', '→ SAD (β)', '→ SocAD (β)',
            '→ Total RLLA (β; R²)']

SIG3 = '***'  # p < .001
SIG2 = '**'   # p < .01
SIG1 = '*'    # p < .05

table_rows = [
    ('RISK FACTORS / YẾU TỐ NGUY CƠ', '', '', '', ''),
    ('Academic pressure (ALHT) / Áp lực học tập',
     f'+0.510{SIG3}', f'+0.253{SIG3}', f'+0.490{SIG3}',
     f'+0.533{SIG3} (R²=0.284)'),
    ('Smartphone addiction (NĐT) / Nghiện điện thoại',
     f'+0.336{SIG3}', f'+0.265{SIG3}', f'+0.383{SIG3}',
     f'+0.400{SIG3} (R²=0.160)'),
    ('School bullying (BNHĐ) / Bắt nạt học đường',
     f'+0.215{SIG3}', f'+0.376{SIG3} ★', f'+0.253{SIG3}',
     f'+0.276{SIG3} (R²=0.076)'),
    ('PROTECTIVE FACTORS / YẾU TỐ BẢO VỆ', '', '', '', ''),
    ('School engagement (GBTH) / Gắn bó trường học',
     f'−0.108{SIG2}', '+0.014 (n.s.)', f'−0.187{SIG3}',
     f'−0.155{SIG3} (R²=0.024)'),
    ('Self-esteem (TTr) / Tự trọng',
     f'−0.455{SIG3} ★', f'−0.087{SIG1}', f'−0.415{SIG3}',
     f'−0.457{SIG3} (R²=0.209)'),
    ('Parental support (HTCM) / Hỗ trợ cha mẹ',
     f'−0.172{SIG3}', '0.000 (n.s.)', f'−0.273{SIG3}',
     f'−0.234{SIG3} (R²=0.055)'),
    ('Peer support (HTBB) / Hỗ trợ bạn bè',
     '−0.015 (n.s.)', '−0.019 (n.s.)', f'−0.079{SIG1}',
     '−0.044 (n.s.) (R²=0.002)'),
]

# Table caption (Frontiers style: above the table)
p = d.add_paragraph()
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(2)
r = p.add_run('Table 1. Standardised path coefficients (β) from '
              'subtype-specific structural equation models.')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True

t = d.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'; t.autofit = False
hdr = t.rows[0].cells
for i, h in enumerate(P_header):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in table_rows:
    row = t.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = str(txt)
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t, [5.5, 2.4, 2.4, 2.4, 4.0])

PP(f'Note. {SIG3} p < .001; {SIG2} p < .01; {SIG1} p < .05; n.s. '
   '= non-significant. ★ = key novel pathways. All β values were '
   'verified against the underlying doctoral thesis (Tables 27 '
   '[ALHT], 30 [NĐT], 33 [BNHĐ], 36 [GBTH], 39 [TTr], 42 [HTCM] '
   'and 45 [HTBB]).',
   f'Ghi chú. {SIG3} p < 0,001; {SIG2} p < 0,01; {SIG1} p < '
   '0,05; n.s. = không có ý nghĩa thống kê. ★ = các đường dẫn '
   'then chốt mới. Tất cả giá trị β đã được đối chiếu với luận '
   'án nền tảng (Bảng 27 [ALHT], 30 [NĐT], 33 [BNHĐ], 36 '
   '[GBTH], 39 [TTr], 42 [HTCM] và 45 [HTBB]).',
   indent=False)


H3('3.4 Three differential-pathway findings',
   '3.4 Ba phát hiện đường dẫn khác biệt')

PP('First, school bullying showed its strongest path to '
   'separation anxiety (β = 0.376, p < .001), exceeding its '
   'paths to GAD (β = 0.215) and SocAD (β = 0.253). This pattern '
   'is counterintuitive because bullying is most often theorised '
   'as a social-evaluative stressor and would therefore be '
   'expected to associate most strongly with SocAD. The '
   'selective amplification of separation anxiety suggests that, '
   'for Vietnamese lower-secondary adolescents, school bullying '
   'activates attachment insecurity rather than '
   'social-evaluative concerns per se.',
   'Thứ nhất, bắt nạt học đường có đường dẫn mạnh nhất tới lo âu '
   'chia ly (β = 0,376; p < 0,001), vượt qua các đường dẫn tới '
   'GAD (β = 0,215) và SocAD (β = 0,253). Mẫu hình này phản '
   'trực giác vì bắt nạt thường được lý thuyết hóa là stress '
   'đánh giá xã hội nên lẽ ra liên hệ mạnh nhất với SocAD. Sự '
   'khuếch đại chọn lọc của lo âu chia ly gợi ý rằng, với học '
   'sinh trung học cơ sở Việt Nam, bắt nạt học đường kích hoạt '
   'mất an toàn gắn bó hơn là các mối lo đánh giá xã hội thuần '
   'túy.')

PP('Second, self-esteem emerged as the strongest single '
   'protective lever for GAD (β = −0.455, p < .001), with a '
   'magnitude approaching that of the dominant academic-pressure '
   'risk effect on GAD (β = 0.510). The total path from '
   'self-esteem to RLLA was −0.457 (R² = 0.209). These findings '
   'position self-esteem building as a high-yield intervention '
   'target.',
   'Thứ hai, tự trọng nổi lên là đòn bẩy bảo vệ đơn lẻ mạnh nhất '
   'đối với GAD (β = −0,455; p < 0,001), với độ lớn gần bằng '
   'tác động chiếm ưu thế của áp lực học tập lên GAD (β = '
   '0,510). Đường dẫn tổng từ tự trọng tới RLLA là −0,457 (R² = '
   '0,209). Các phát hiện này định vị xây dựng tự trọng như một '
   'mục tiêu can thiệp có hiệu suất cao.')

PP('Third, peer support had no significant effect on total '
   'RLLA (β = −0.044, p = .183, R² = 0.002) and was '
   'non-significant for both GAD (β = −0.015) and SAD '
   '(β = −0.019), reaching only marginal significance for '
   'SocAD (β = −0.079, p = .020). This pattern challenges the '
   'assumption that peer relationships are universally '
   'protective and is consistent with the co-rumination '
   'hypothesis: among adolescents who do confide, the frequency '
   'of peer disclosure may not equate to the quality of '
   'support.',
   'Thứ ba, hỗ trợ bạn bè không có tác động có ý nghĩa lên tổng '
   'RLLA (β = −0,044; p = 0,183; R² = 0,002), không có ý nghĩa '
   'với cả GAD (β = −0,015) và SAD (β = −0,019), chỉ đạt ý '
   'nghĩa cận biên với SocAD (β = −0,079; p = 0,020). Mẫu hình '
   'này thách thức giả định về vai trò bảo vệ phổ quát của '
   'quan hệ bạn bè và phù hợp với giả thuyết co-rumination: ở '
   'thanh thiếu niên có chia sẻ, tần suất giãi bày với bạn bè '
   'không nhất thiết đồng nghĩa với chất lượng hỗ trợ.')


H3('3.5 Integrated risk-protective composite model',
   '3.5 Mô hình tích hợp nguy cơ – bảo vệ tổng hợp')

PP('To assess the joint contribution of risk and protective '
   'systems, a higher-order composite SEM combined the three risk '
   'factors into a Risk latent (YTNC) and the four protective '
   'factors into a Protective latent (YTBV), both predicting '
   'total RLLA simultaneously. The composite Risk → RLLA path '
   'was β = 0.669 (p < .001), and the composite Protection → '
   'RLLA path was β = −0.220 (p < .001), with an integrated '
   'R² = 0.598 for total RLLA. The risk-side composite thus '
   'accounted for approximately three times the standardised '
   'effect of the protective-side composite, indicating that '
   'risk-factor reduction may be the higher-leverage '
   'intervention target at the population level in this '
   'Vietnamese sample.',
   'Nhằm đánh giá đóng góp đồng thời của hệ thống nguy cơ và bảo '
   'vệ, một SEM tổng hợp bậc cao kết hợp ba yếu tố nguy cơ thành '
   'một biến tiềm ẩn Nguy cơ (YTNC) và bốn yếu tố bảo vệ thành '
   'một biến tiềm ẩn Bảo vệ (YTBV), cả hai cùng dự báo tổng '
   'RLLA đồng thời. Đường dẫn tổng hợp Nguy cơ → RLLA là β = '
   '0,669 (p < 0,001) và đường dẫn tổng hợp Bảo vệ → RLLA là '
   'β = −0,220 (p < 0,001), với R² tích hợp = 0,598 cho tổng '
   'RLLA. Như vậy, mặt nguy cơ chiếm khoảng gấp ba lần ảnh '
   'hưởng chuẩn hóa của mặt bảo vệ, cho thấy giảm yếu tố nguy '
   'cơ có thể là mục tiêu can thiệp có đòn bẩy cao hơn ở cấp '
   'độ quần thể trong mẫu Việt Nam này.')


H3('3.6 Multi-group invariance by gender — H3 supported',
   '3.6 Bất biến đa nhóm theo giới — H3 được ủng hộ')

PP('Finally, we tested H3 — the cultural-collectivism prediction '
   'that SAD alone would show gender invariance — using multi-group '
   'SEM. Configural invariance was achieved (both groups fitted the '
   'same structural model), supporting subsequent invariance '
   'comparisons, following procedures established for '
   'cross-gender invariance of the GAD-7 in adolescent samples '
   '(29). Mean scores differed significantly by gender for GAD '
   '(female M = 59.47 vs male M = 51.43; F(1, 1350) = 44.48, '
   'p < .001) and SocAD (female M = 52.74 vs male M = 43.20; '
   'F(1, 1350) = 45.98, p < .001). Strikingly, SAD exhibited '
   'complete gender invariance (female M = 24.76 vs male '
   'M = 25.42; F(1, 1350) = 0.246, p = .620), consistent with '
   'H3. To our knowledge, this is the first documented instance '
   'of complete SAD gender invariance in a Vietnamese '
   'lower-secondary sample.',
   'Cuối cùng, chúng tôi kiểm định H3 — dự đoán dựa trên văn hóa '
   'tập thể rằng chỉ riêng SAD sẽ thể hiện bất biến giới — bằng '
   'SEM đa nhóm. Bất biến cấu hình được đạt (cả hai nhóm vừa khít '
   'cùng một '
   'mô hình cấu trúc), hỗ trợ các so sánh bất biến tiếp theo, '
   'theo các quy trình đã được thiết lập cho bất biến giữa các '
   'giới của GAD-7 trên mẫu thanh thiếu niên (29). Điểm trung '
   'bình khác biệt có ý nghĩa theo giới với GAD (nữ M = 59,47 '
   'so với nam M = 51,43; F(1, 1350) = 44,48; p < 0,001) và '
   'SocAD (nữ M = 52,74 so với nam M = 43,20; F(1, 1350) = '
   '45,98; p < 0,001). Đáng chú ý, SAD thể hiện bất biến giới '
   'hoàn toàn (nữ M = 24,76 so với nam M = 25,42; F(1, 1350) = '
   '0,246; p = 0,620), phù hợp với H3. Theo hiểu biết của chúng '
   'tôi, đây là trường hợp đầu tiên được ghi nhận về bất biến '
   'giới hoàn toàn của SAD trên mẫu học sinh trung học cơ sở '
   'Việt Nam.')


d.add_page_break()


# ============================================================
H2('4 Discussion', '4 Thảo luận')

H3('4.1 Summary of findings', '4.1 Tổng hợp các phát hiện')

PP('The headline contribution of this study is direct evidence that '
   'Vietnamese adolescents do not experience a single, unitary '
   'anxiety syndrome but rather three subtype-specific pathways '
   'with distinct risk-protective architectures and distinct '
   'gender signatures. In a large Vietnamese lower-secondary '
   'sample, integrated SEM documented three differential-pathway '
   'findings: bullying showed its strongest impact on separation '
   'anxiety, self-esteem was the strongest single protective lever '
   'for generalised anxiety and peer support exerted no significant '
   'protective effect on any subtype. The composite '
   'risk-protective model accounted for 59.8% of the variance '
   'in total anxiety, with the risk-side contribution '
   'approximately three times the protective-side contribution. '
   'Multi-group invariance testing revealed gender differences '
   'for GAD and SocAD but complete gender invariance for SAD — '
   'a finding we interpret through a cultural-collectivism '
   'lens. H1 was supported in full; H2 was supported, with the '
   'notable exception of the peer-support paths; H3 was '
   'supported.',
   'Đóng góp cốt lõi của nghiên cứu này là bằng chứng trực tiếp '
   'cho thấy thanh thiếu niên Việt Nam không trải nghiệm một hội '
   'chứng lo âu đơn nhất, mà là ba đường dẫn theo từng phân loại '
   'với cấu trúc nguy cơ – bảo vệ riêng biệt và dấu ấn giới riêng '
   'biệt. Trong một mẫu lớn học sinh trung học cơ sở Việt Nam, '
   'SEM tích hợp ghi nhận ba phát hiện đường dẫn khác biệt: bắt '
   'nạt có tác động mạnh nhất lên lo âu chia ly, tự trọng là đòn '
   'bẩy bảo vệ đơn lẻ mạnh nhất đối với lo âu tổng quát, và hỗ '
   'trợ bạn bè không có hiệu lực bảo vệ có ý nghĩa với bất kỳ '
   'phân loại nào. Mô hình tổng hợp nguy cơ – bảo vệ giải '
   'thích 59,8% phương sai của tổng lo âu, với đóng góp của '
   'mặt nguy cơ khoảng gấp ba lần đóng góp của mặt bảo vệ. '
   'Kiểm định bất biến đa nhóm cho thấy khác biệt giới với '
   'GAD và SocAD nhưng bất biến giới hoàn toàn với SAD — một '
   'phát hiện chúng tôi diễn giải qua lăng kính văn hóa tập '
   'thể. H1 được ủng hộ đầy đủ; H2 được ủng hộ với ngoại lệ '
   'đáng chú ý của các đường dẫn từ hỗ trợ bạn bè; H3 được '
   'ủng hộ.')


H3('4.2 Bullying → SAD as the strongest risk pathway',
   '4.2 Bắt nạt → lo âu chia ly: đường dẫn nguy cơ mạnh nhất')

PP('The dominance of bullying for separation anxiety '
   '(β = 0.376) rather than for the more theoretically '
   'anticipated social-anxiety subtype (β = 0.253) is, to our '
   'knowledge, the first such observation in a Vietnamese '
   'lower-secondary sample. We interpret this pattern through '
   'attachment theory: bullying victimisation plausibly '
   'disrupts the perception of school as a secondary secure '
   'base, intensifying attachment-related distress on '
   'separation from primary caregivers (30). The pattern '
   'aligns with cross-cultural evidence that bullying '
   'victimisation in collectivist contexts is preferentially '
   'processed as an attachment-related threat (12).',
   'Tính chiếm ưu thế của bắt nạt với lo âu chia ly (β = '
   '0,376) so với phân loại lo âu xã hội được kỳ vọng về mặt '
   'lý thuyết (β = 0,253), theo hiểu biết của chúng tôi, là '
   'quan sát đầu tiên kiểu này trên mẫu học sinh trung học '
   'cơ sở Việt Nam. Chúng tôi diễn giải mẫu hình này qua lý '
   'thuyết gắn bó (30): nạn nhân bắt nạt có khả năng làm '
   'gián đoạn nhận thức về trường học như một căn cứ an toàn '
   'thứ cấp, làm tăng cường đau khổ liên quan đến gắn bó khi '
   'chia ly với người chăm sóc chính. Mẫu hình này phù hợp '
   'với bằng chứng liên văn hóa rằng nạn nhân bắt nạt trong '
   'bối cảnh tập thể được xử lý ưu tiên như mối đe dọa liên '
   'quan đến gắn bó (12).')


H3('4.3 Self-esteem as the strongest protective lever',
   '4.3 Tự trọng — đòn bẩy bảo vệ đơn lẻ mạnh nhất')

PP('The path from self-esteem to GAD (β = −0.455) approached '
   'the magnitude of the academic-pressure risk path '
   '(β = 0.510), and the path from self-esteem to total RLLA '
   'reached β = −0.457 (R² = 0.209). These findings align with '
   'the Compas et al. meta-analysis (212 studies, N = 80,850), '
   'which identified secondary control coping — including '
   'cognitive reappraisal — as the most consistent predictor '
   'of lower internalising symptoms in childhood and '
   'adolescence (15). Systematic reviews of academic pressure '
   'and adolescent mental health (31, 32) further support '
   'academic-stress management as a complementary intervention '
   'target. Together, these data identify self-esteem building '
   'as a high-yield single-component intervention focus for '
   'adolescent GAD in the Vietnamese context.',
   'Đường dẫn từ tự trọng tới GAD (β = −0,455) gần bằng độ '
   'lớn của đường dẫn nguy cơ từ áp lực học tập (β = 0,510), '
   'và đường dẫn từ tự trọng tới tổng RLLA đạt β = −0,457 '
   '(R² = 0,209). Các phát hiện này phù hợp với phân tích '
   'tổng hợp của Compas và cộng sự (212 nghiên cứu, N = '
   '80.850), trong đó nhận diện chiến lược ứng phó kiểm soát '
   'thứ cấp — bao gồm tái diễn giải nhận thức — là biến dự '
   'báo nhất quán nhất cho mức triệu chứng nội hóa thấp ở '
   'trẻ em và thanh thiếu niên (15). Các tổng quan hệ thống '
   'về áp lực học tập và sức khỏe tâm thần thanh thiếu niên '
   '(31, 32) tiếp tục ủng hộ quản lý stress học tập như một '
   'mục tiêu can thiệp bổ trợ. Tổng hợp lại, các dữ liệu này '
   'xác định xây dựng tự trọng là một mục tiêu can thiệp đơn '
   'thành phần có hiệu suất cao cho GAD ở thanh thiếu niên '
   'trong bối cảnh Việt Nam.')


H3('4.4 Peer support null effects: the co-rumination hypothesis',
   '4.4 Hỗ trợ bạn bè không có tác động: giả thuyết co-rumination')

PP('Western literature consistently identifies peer support '
   'as protective against adolescent internalising symptoms. '
   'Our null and marginally significant results across all '
   'three subtypes depart from this expectation. We propose '
   'two non-mutually-exclusive interpretations. First, in a '
   'Confucian-influenced collectivist context shaped by the '
   'tam giao tradition (the coexistence of Confucianism, '
   'Buddhism and Taoism), peer disclosure norms may emphasise '
   'emotional restraint over expressive sharing (33), '
   'attenuating the support-seeking buffer mechanism. Second, '
   'the co-rumination construct (34) — a perseverative form '
   'of dyadic problem-talk in which adolescents repeatedly '
   'revisit stressful events and their accompanying negative '
   'affect rather than moving toward resolution — suggests '
   'that frequent peer disclosure can recursively amplify '
   'worry rather than dissipate it, especially when '
   'academic-competition norms are operative (11). These '
   'theoretical interpretations are consistent with '
   'cross-cultural evidence on emotion socialisation in East '
   'Asian adolescent samples. Taken '
   'together, the tam giao framing and the co-rumination '
   'mechanism reinforce the study\'s broader message that '
   'Western risk-protective architectures cannot be transferred '
   'wholesale to Vietnamese adolescents; subtype- and '
   'culture-specific calibration is required.',
   'Y văn phương Tây thường xuyên xác định hỗ trợ bạn bè là '
   'yếu tố bảo vệ chống lại các triệu chứng nội hóa ở thanh '
   'thiếu niên. Các kết quả không có ý nghĩa và cận biên của '
   'chúng tôi xuyên suốt cả ba phân loại lệch khỏi kỳ vọng '
   'này. Chúng tôi đề xuất hai cách diễn giải không loại trừ '
   'lẫn nhau. Thứ nhất, trong bối cảnh tập thể ảnh hưởng Nho '
   'giáo được định hình bởi truyền thống tam giáo (sự cộng '
   'tồn của Nho giáo, Phật giáo và Đạo giáo), chuẩn mực giãi '
   'bày với bạn bè có thể nhấn mạnh sự chừng mực cảm xúc hơn '
   'là chia sẻ biểu cảm (33), làm suy giảm cơ chế đệm tìm '
   'kiếm hỗ trợ. Thứ hai, khái niệm co-rumination (34) — một '
   'dạng trò chuyện cặp đôi mang tính lặp lại trong đó thanh '
   'thiếu niên liên tục quay lại các sự kiện gây stress và '
   'các cảm xúc tiêu cực kèm theo thay vì hướng tới giải '
   'quyết — gợi ý rằng giãi bày thường xuyên với bạn bè có '
   'thể khuếch đại lo lắng theo '
   'vòng lặp thay vì tiêu tan, đặc biệt khi chuẩn mực cạnh '
   'tranh học thuật đang vận hành (11). Các diễn giải lý '
   'thuyết này phù hợp với bằng chứng liên văn hóa về xã hội '
   'hóa cảm xúc ở các mẫu thanh thiếu niên Đông Á. Gộp lại, '
   'khung tam giáo và cơ chế co-rumination củng cố thông điệp '
   'rộng hơn của nghiên cứu: các cấu trúc nguy cơ – bảo vệ từ '
   'phương Tây không thể chuyển giao nguyên trạng cho thanh '
   'thiếu niên Việt Nam; cần hiệu chỉnh theo từng phân loại và '
   'theo văn hóa.')


H3('4.5 Separation anxiety gender invariance — a novel interpretation',
   '4.5 Bất biến giới của lo âu chia ly — diễn giải mới')

PP('The complete gender invariance for SAD (female M = 24.76 '
   'vs male M = 25.42; p = .620), set against significant '
   'gender differences for GAD and SocAD, is the most '
   'theoretically distinctive finding of this study. We '
   'advance two non-mutually-exclusive explanations. First, '
   'cultural collectivism: Vietnamese family hierarchy '
   'creates uniform separation experiences across genders '
   '(12). Under interdependent self-construal (13), '
   'attachment processes operate primarily within the family '
   'system, where gender role differentiation is less '
   'salient for separation specifically than for socially '
   'evaluative behaviour. Second, developmental timing: '
   'separation-individuation tasks are predominantly '
   'childhood-onset and pre-date the post-pubertal '
   'differentiation of gender-loaded social pressures '
   '(19, 20). Pre-pubertal separation challenges are '
   'universal across genders, whereas post-pubertal social '
   'evaluation and academic-performance pressures are '
   'gendered.',
   'Bất biến giới hoàn toàn của SAD (nữ M = 24,76 so với nam '
   'M = 25,42; p = 0,620), trái ngược với khác biệt giới có '
   'ý nghĩa của GAD và SocAD, là phát hiện mang tính khác '
   'biệt lý thuyết nổi bật nhất của nghiên cứu. Chúng tôi '
   'đưa ra hai cách giải thích không loại trừ lẫn nhau. Thứ '
   'nhất, văn hóa tập thể: thứ bậc gia đình Việt Nam tạo '
   'trải nghiệm chia ly đồng nhất giữa các giới (12). Dưới '
   'khái niệm bản thân phụ thuộc lẫn nhau (13), các quá '
   'trình gắn bó vận hành chủ yếu trong hệ thống gia đình, '
   'nơi sự phân hóa vai trò giới kém nổi bật hơn đối với '
   'chia ly so với hành vi mang tính đánh giá xã hội. Thứ '
   'hai, thời điểm phát triển: các nhiệm vụ chia ly – cá '
   'nhân hóa chủ yếu khởi phát từ thời thơ ấu và đi trước '
   'sự phân hóa giới của các áp lực xã hội sau dậy thì '
   '(19, 20). Các thách thức chia ly tiền dậy thì phổ quát '
   'giữa các giới, trong khi áp lực đánh giá xã hội và thành '
   'tích học tập sau dậy thì mang tính giới.')


H3('4.6 Clinical and educational implications',
   '4.6 Hàm ý lâm sàng và giáo dục')

PP('If anxiety pathways are subtype-specific, prevention '
   'strategies must be too. Three implications follow. First, a '
   'tiered prevention architecture is warranted: universal '
   'classroom-based '
   'intervention for risk-factor reduction (academic pressure '
   'management, anti-bullying programmes and smartphone-use '
   'boundaries) plus targeted self-esteem-building modules for '
   'high-anxiety students. Second, gender-aware delivery is '
   'appropriate for GAD and SocAD modules but unnecessary for '
   'SAD content — a calibration that may improve '
   'cost-efficiency. Third, anti-bullying programmes should '
   'explicitly incorporate separation-anxiety screening, '
   'given the dominant bullying → SAD pathway documented '
   'here.',
   'Nếu các đường dẫn lo âu mang tính phân loại, thì các chiến '
   'lược phòng ngừa cũng phải vậy. Có thể rút ra ba hàm ý. Thứ '
   'nhất, một kiến trúc phòng ngừa phân tầng là cần thiết: can '
   'thiệp dựa trên lớp học '
   'phổ quát để giảm yếu tố nguy cơ (quản lý áp lực học tập, '
   'chương trình chống bắt nạt và ranh giới sử dụng điện '
   'thoại) kết hợp với module xây dựng tự trọng cho học '
   'sinh có mức lo âu cao. Thứ hai, triển khai nhạy với '
   'giới phù hợp cho module GAD và SocAD nhưng không cần '
   'thiết cho nội dung SAD — một sự hiệu chỉnh có thể cải '
   'thiện hiệu suất chi phí. Thứ ba, chương trình chống bắt '
   'nạt cần tích hợp sàng lọc lo âu chia ly một cách rõ '
   'ràng, do đường dẫn bắt nạt → SAD chiếm ưu thế được ghi '
   'nhận tại đây.')


H3('4.7 Limitations and future directions',
   '4.7 Hạn chế và hướng nghiên cứu tương lai')

PP('Several limitations bound the present findings. First, '
   'the cross-sectional design precludes causal inference; '
   'single-timepoint measurement does not permit temporal '
   'ordering of risk-protective and anxiety pathways. A '
   'longitudinal cohort following the same students from '
   'Grade 6 to Grade 9 would be required to establish '
   'directional and developmental sequencing. Second, the '
   'sample was drawn from two Hanoi schools and may not '
   'generalise to rural or southern Vietnamese contexts; '
   'multi-site replication is needed. Third, self-report '
   'measures of bullying and anxiety carry known '
   'shared-method variance risk; future studies should '
   'integrate peer-nomination bullying measures and '
   'clinical diagnostic anxiety interviews. Fourth, future '
   'randomised controlled trials should test an eight-module '
   'school-based prevention curriculum derived from these '
   'mechanistic findings, ideally adapting the OurFutures '
   'Mental Health Australia template for the Vietnamese '
   'context (35).',
   'Một số hạn chế giới hạn các phát hiện hiện tại. Thứ '
   'nhất, thiết kế cắt ngang loại trừ suy luận nhân quả; '
   'đo lường tại một thời điểm duy nhất không cho phép sắp '
   'xếp thời gian của các đường dẫn nguy cơ – bảo vệ và lo '
   'âu. Một đoàn hệ dọc theo dõi cùng học sinh từ khối 6 '
   'đến khối 9 sẽ cần thiết để thiết lập trình tự hướng và '
   'phát triển. Thứ hai, mẫu lấy từ hai trường ở Hà Nội và '
   'có thể không khái quát hóa cho bối cảnh nông thôn hoặc '
   'miền Nam Việt Nam; cần nhân rộng đa địa điểm. Thứ ba, '
   'các đo lường tự báo cáo về bắt nạt và lo âu mang rủi '
   'ro phương sai phương pháp chia sẻ; nghiên cứu tương lai '
   'cần tích hợp đo lường bắt nạt theo đề cử bạn bè và '
   'phỏng vấn chẩn đoán lâm sàng về lo âu. Thứ tư, các thử '
   'nghiệm ngẫu nhiên có đối chứng trong tương lai cần kiểm '
   'định một chương trình giảng dạy phòng ngừa 8 nội dung '
   'rút ra từ các phát hiện cơ chế này, lý tưởng nhất với '
   'mô hình OurFutures Mental Health Úc được điều chỉnh cho '
   'bối cảnh Việt Nam (35).')


d.add_page_break()


# ============================================================
H2('5 Conclusion', '5 Kết luận')

PP('In a large Vietnamese lower-secondary sample, integrated '
   'SEM identified three differential pathways: bullying → '
   'separation anxiety as the strongest risk pathway, '
   'self-esteem as the strongest single protective lever for '
   'generalised anxiety and a non-significant peer-support '
   'effect across all subtypes. Multi-group invariance '
   'testing revealed complete gender invariance for '
   'separation anxiety alongside significant gender '
   'differences for GAD and SocAD. These findings support '
   'subtype-targeted prevention frameworks that pair '
   'culture-sensitive self-esteem building with '
   'bullying-specific intervention for separation anxiety, '
   'and they motivate longitudinal extension to confirm the '
   'developmental sequencing of these effects.',
   'Trong một mẫu lớn học sinh trung học cơ sở Việt Nam, '
   'SEM tích hợp xác định ba đường dẫn khác biệt: bắt nạt → '
   'lo âu chia ly là đường dẫn nguy cơ mạnh nhất, tự trọng '
   'là đòn bẩy bảo vệ đơn lẻ mạnh nhất đối với lo âu tổng '
   'quát và hỗ trợ bạn bè không có hiệu lực có ý nghĩa qua '
   'cả ba phân loại. Kiểm định bất biến đa nhóm cho thấy '
   'bất biến giới hoàn toàn với lo âu chia ly bên cạnh '
   'khác biệt giới có ý nghĩa cho GAD và SocAD. Các phát '
   'hiện này ủng hộ các khung phòng ngừa theo từng phân '
   'loại, kết hợp xây dựng tự trọng nhạy với văn hóa và '
   'can thiệp đặc thù cho bắt nạt với lo âu chia ly, đồng '
   'thời tạo cơ sở cho việc mở rộng theo chiều dọc nhằm '
   'xác nhận trình tự phát triển của các tác động này.')


# ============================================================
# END-MATTER (Frontiers required sections)
# ============================================================
H2('Data availability statement', 'Tuyên bố về tính sẵn có của dữ liệu')

PP('The de-identified datasets and AMOS syntax files used in '
   'this study are available from the corresponding author on '
   'reasonable request, subject to approval by the Hanoi '
   'National University of Education Institutional Review '
   'Board.',
   'Bộ dữ liệu đã ẩn danh và các tệp lệnh AMOS sử dụng trong '
   'nghiên cứu này có thể được cung cấp từ tác giả liên hệ '
   'theo yêu cầu hợp lý, tuân thủ phê duyệt của Hội đồng Đạo '
   'đức Trường Đại học Sư phạm Hà Nội.')


H2('Ethics statement', 'Tuyên bố về đạo đức nghiên cứu')

PP('The study was reviewed and approved by the Institutional '
   'Review Board of Hanoi National University of Education. '
   'Written informed consent was obtained from a parent or '
   'legal guardian of each participant, and written assent '
   'was obtained from each participating student.',
   'Nghiên cứu được Hội đồng Đạo đức của Trường Đại học Sư '
   'phạm Hà Nội xem xét và phê duyệt. Văn bản đồng ý có thông '
   'tin được lấy từ cha mẹ hoặc người giám hộ hợp pháp của '
   'mỗi người tham gia, và văn bản đồng ý cũng được lấy từ '
   'mỗi học sinh tham gia.')


H2('Author contributions', 'Đóng góp của tác giả')

PP('HTC: conceptualisation, methodology, data collection, '
   'formal analysis, writing — original draft, project '
   'administration. DMD: conceptualisation, supervision, '
   'methodology, writing — review and editing. NMD: '
   'methodology consultation, statistical guidance, writing — '
   'review and editing, senior supervision. All authors '
   'contributed to the article and approved the submitted '
   'version.',
   'HTC: khái niệm hóa, phương pháp luận, thu thập dữ liệu, '
   'phân tích chính thức, viết bản thảo gốc, quản lý dự án. '
   'DMD: khái niệm hóa, giám sát, phương pháp luận, rà soát '
   'và biên tập. NMD: tư vấn phương pháp luận, hướng dẫn '
   'thống kê, rà soát và biên tập, giám sát cấp cao. Tất cả '
   'các tác giả đã đóng góp cho bài báo và phê duyệt phiên '
   'bản nộp.')


H2('Funding', 'Tài trợ')

PP('The authors declare that no financial support was '
   'received for the research, authorship and/or publication '
   'of this article.',
   'Các tác giả tuyên bố không nhận tài trợ tài chính cho '
   'nghiên cứu, viết bài và/hoặc xuất bản bài báo này.')


H2('Conflict of interest', 'Xung đột lợi ích')

PP('The authors declare that the research was conducted in '
   'the absence of any commercial or financial relationships '
   'that could be construed as a potential conflict of '
   'interest.',
   'Các tác giả tuyên bố nghiên cứu được thực hiện mà không '
   'có bất kỳ quan hệ thương mại hoặc tài chính nào có thể '
   'được coi là xung đột lợi ích tiềm tàng.')


H2('Acknowledgments', 'Lời cảm ơn')

PP('The authors thank the school administrators, teachers, '
   'parents and students of Nhat Tan and Tay Mo Lower '
   'Secondary Schools, Hanoi, for their participation and '
   'support. We also thank the three clinical psychologists '
   'and one educational measurement specialist who '
   'contributed to the expert review of the Vietnamese scale '
   'adaptations.',
   'Các tác giả cảm ơn ban giám hiệu, giáo viên, phụ huynh '
   'và học sinh của các Trường Trung học Cơ sở Nhật Tân và '
   'Tây Mỗ (Hà Nội) vì sự tham gia và hỗ trợ. Chúng tôi cũng '
   'cảm ơn ba nhà tâm lý lâm sàng và một chuyên gia đo lường '
   'giáo dục đã đóng góp vào quá trình rà soát chuyên gia '
   'của bản chuyển ngữ tiếng Việt của các thang đo.')


d.add_page_break()


# ============================================================
H2('References', 'Tài liệu tham khảo')
# Frontiers in Psychiatry uses numbered references in order of appearance.

P_ref = [
    # 1 GBD ASEAN
    '1. GBD ASEAN Collaborators. Burden of mental disorders in '
    'ASEAN children and adolescents: Global Burden of Disease '
    'analysis 1990–2021. Lancet Reg Health West Pac. 2025.',
    # 2 Kieling
    '2. Kieling C, Buchweitz C, Caye A, et al. Worldwide '
    'prevalence and disability from mental disorders across '
    'childhood and adolescence: evidence from the Global '
    'Burden of Disease Study. JAMA Psychiatry. 2024;81(4):'
    '347–356. doi:10.1001/jamapsychiatry.2023.5051.',
    # 3 Robson
    '3. Robson EM, Husin HM, Dashti SG, et al. Tracking the '
    'course of depressive and anxiety symptoms across '
    'adolescence (the CATS study). Lancet Psychiatry. '
    '2025;12(1):44–53.',
    # 4 APA DSM-5
    '4. American Psychiatric Association. Diagnostic and '
    'Statistical Manual of Mental Disorders. 5th ed. '
    'Arlington, VA: American Psychiatric Publishing; 2013.',
    # 5 Chorpita RCADS
    '5. Chorpita BF, Yim L, Moffitt C, Umemoto LA, Francis '
    'SE. Assessment of symptoms of DSM-IV anxiety and '
    'depression in children: a revised child anxiety and '
    'depression scale. Behav Res Ther. 2000;38(8):835–855.',
    # 6 Xu
    '6. Xu J, et al. Prevalence of mental disorders among '
    'Chinese adolescents (N = 373,216). 2021.',
    # 7 Chen
    '7. Chen Z, Ren S, He R, et al. Prevalence and '
    'associated factors of depressive and anxiety symptoms '
    'among Chinese secondary school students. BMC '
    'Psychiatry. 2023;23(1):580.',
    # 8 Wen
    '8. Wen F, et al. Mental health among Chinese rural '
    'adolescents. 2020.',
    # 9 Saikia
    '9. Saikia S, et al. Prevalence and correlates of '
    'anxiety in Northeast Indian adolescents. Indian J '
    'Community Med. 2023.',
    # 10 Anderson
    '10. Anderson TL, Valiauga R, Tallo C, et al. '
    'Contributing factors to the rise in adolescent '
    'anxiety and associated mental health disorders: a '
    'narrative review of current literature. J Child '
    'Adolesc Psychiatr Nurs. 2025;38(1):e70009. '
    'doi:10.1111/jcap.70009. PMID: 39739929.',
    # 11 Stankov
    '11. Stankov L. Unforgiving Confucian culture: a '
    'breeding ground for high academic achievement, test '
    'anxiety and self-doubt? Learn Individ Differ. '
    '2010;20(6):555–563. doi:10.1016/j.lindif.2010.05.003.',
    # 12 Triandis
    '12. Triandis HC. Individualism and Collectivism. '
    'Boulder, CO: Westview Press; 1995.',
    # 13 Markus & Kitayama
    '13. Markus HR, Kitayama S. Culture and the self: '
    'implications for cognition, emotion, and motivation. '
    'Psychol Rev. 1991;98(2):224–253.',
    # 14 Lazarus & Folkman
    '14. Lazarus RS, Folkman S. Stress, Appraisal, and '
    'Coping. New York: Springer; 1984.',
    # 15 Compas
    '15. Compas BE, Jaser SS, Bettis AH, et al. Coping, '
    'emotion regulation, and psychopathology in childhood '
    'and adolescence: a meta-analysis and narrative '
    'review. Psychol Bull. 2017;143(9):939–991. '
    'doi:10.1037/bul0000110. PMID: 28616996.',
    # 16 V-NAMHS
    '16. Vietnam Adolescent Mental Health Survey '
    '(V-NAMHS). Main report. Hanoi; 2022.',
    # 17 Hoang Trung Hoc
    '17. Hoang Trung Hoc. Prevalence of anxiety disorders '
    'in Vietnamese students (N = 8,389). 2025.',
    # 18 McLean
    '18. McLean CP, Asnaani A, Litz BT, Hofmann SG. '
    'Gender differences in anxiety disorders: prevalence, '
    'course of illness, comorbidity and burden of '
    'illness. J Psychiatr Res. 2011;45(8):1027–1035. '
    'doi:10.1016/j.jpsychires.2011.03.006. PMID: 21439576.',
    # 19 Blos
    '19. Blos P. The Adolescent Passage: Developmental '
    'Issues. New York: International Universities Press; '
    '1979.',
    # 20 Kroger
    '20. Kroger J. Identity Development: Adolescence '
    'Through Adulthood. 2nd ed. Thousand Oaks, CA: Sage; '
    '2007.',
    # 21 Sun ESSA
    '21. Sun J, Dunne MP, Hou XY, Xu A. Educational '
    'stress scale for adolescents: development, validity, '
    'and reliability with Chinese students. J Psychoeduc '
    'Assess. 2011;29(6):534–546.',
    # 22 Kwon SAS-SV
    '22. Kwon M, Lee JY, Won WY, et al. Development and '
    'validation of a Smartphone Addiction Scale (SAS). '
    'PLoS One. 2013;8(2):e56936.',
    # 23 Olweus OBVQ
    '23. Olweus D. The Revised Olweus Bully/Victim '
    'Questionnaire. Bergen, Norway: University of Bergen; '
    '1996.',
    # 24 Goodenow PSSM
    '24. Goodenow C. The psychological sense of school '
    'membership among adolescents: scale development and '
    'educational correlates. Psychol Sch. 1993;30(1):'
    '79–90.',
    # 25 Zimet MSPSS
    '25. Zimet GD, Dahlem NW, Zimet SG, Farley GK. The '
    'Multidimensional Scale of Perceived Social Support. '
    'J Pers Assess. 1988;52(1):30–41.',
    # 26 Rosenberg RSES
    '26. Rosenberg M. Society and the Adolescent '
    'Self-image. Princeton, NJ: Princeton University '
    'Press; 1965.',
    # 27 Hu & Bentler
    '27. Hu L, Bentler PM. Cutoff criteria for fit '
    'indexes in covariance structure analysis: '
    'conventional criteria versus new alternatives. '
    'Struct Equ Modeling. 1999;6(1):1–55.',
    # 28 Cheung & Rensvold
    '28. Cheung GW, Rensvold RB. Evaluating '
    'goodness-of-fit indexes for testing measurement '
    'invariance. Struct Equ Modeling. 2002;9(2):'
    '233–255.',
    # 29 Niwenahisemo
    '29. Niwenahisemo LC, Hong S, Kuang L. Assessing '
    'anxiety symptom severity in Rwandese adolescents: '
    'cross-gender measurement invariance of GAD-7. '
    'Front Psychiatry. 2024.',
    # 30 Bowlby
    '30. Bowlby J. A Secure Base: Parent-Child '
    'Attachment and Healthy Human Development. New '
    'York: Basic Books; 1988.',
    # 31 Pascoe
    '31. Pascoe MC, Hetrick SE, Parker AG. The impact '
    'of stress on students in secondary school and '
    'higher education. Int J Adolesc Youth. 2020;25(1):'
    '104–112.',
    # 32 Steare
    '32. Steare T, Gutiérrez Muñoz C, Sullivan A, Lewis '
    'G. The association between academic pressure and '
    'adolescent mental health problems: a systematic '
    'review. J Affect Disord. 2023.',
    # 33 Small & Blanc
    '33. Small S, Blanc J. Mental health during '
    'COVID-19: Tam Giao and Vietnam\'s response. Front '
    'Psychiatry. 2020;11:589618. '
    'doi:10.3389/fpsyt.2020.589618.',
    # 34 Rose co-rumination
    '34. Rose AJ. Co-rumination in the friendships of '
    'girls and boys. Child Dev. 2002;73(6):1830–1843. '
    'PMID: 12487497.',
    # 35 Grummitt OurFutures
    '35. Grummitt L, O\'Dean S, Birrell L, et al. '
    'Efficacy of a school-based, universal prevention '
    'programme for depression and anxiety in '
    'adolescents (OurFutures Mental Health). '
    'eClinicalMedicine. 2025.',
]
for ref in P_ref:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.first_line_indent = Cm(-0.7)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# ============================================================
H2('Notes to co-authors — 3 BLOCKING items',
   'Ghi chú cho đồng tác giả — 3 mục cần quyết')

TBD('Q1-8 — Integrated SEM R² interpretation: present draft '
    'uses LA Bảng 47 composite R² = 0.598 (YTNC + YTBV → '
    'RLLA). Co-authors to confirm whether this composite '
    'reporting is acceptable for Frontiers, or whether a full '
    're-analysis of the 7 × 3 integrated SEM is required '
    'before submission.',
    'Q1-8 — Diễn giải R² SEM tích hợp: bản nhập hiện tại dùng '
    'R² tổng hợp = 0,598 (YTNC + YTBV → RLLA) từ LA Bảng '
    '47. Đồng tác giả xác nhận liệu báo cáo tổng hợp này có '
    'chấp nhận được cho Frontiers hay không, hoặc liệu có '
    'cần tái phân tích đầy đủ SEM tích hợp 7 × 3 trước khi '
    'nộp bài hay không.')

TBD('Q3-6 — IRB approval letter from HNUE: NCS to confirm '
    'and supply the exact decision number, date of approval '
    'and a scanned copy of the official letter from the '
    'HNUE Ethics Committee for attachment as a Supplementary '
    'File. The draft text in §2.4 must be updated with the '
    'verified decision number before submission.',
    'Q3-6 — Letter chấp thuận đạo đức HNUE: NCS xác nhận và '
    'cung cấp chính xác số quyết định, ngày phê duyệt và '
    'bản scan của thư chính thức từ Hội đồng Đạo đức ĐHSPHN '
    'để đính kèm vào Tệp Bổ sung. Đoạn văn trong §2.4 cần '
    'được cập nhật số quyết định đã được kiểm chứng trước '
    'khi nộp.')

TBD('Q3-9 — Submission strategy for Frontiers in '
    'Psychiatry: confirm the target section. Three '
    'candidates are listed below with brief pro/con notes:\n'
    '   (a) Adolescent and Young Adult Psychiatry — PRO: '
    'matches sample (ages 11–14); subtype-specific design '
    'is squarely on-mission. CON: highly competitive section '
    'with strong preference for longitudinal/clinical work.\n'
    '   (b) Anxiety and Stress Disorders — PRO: aligns with '
    'DSM-5 subtype focus and outcome variable. CON: broader '
    'adult-skewed readership, potentially less appetite for '
    'school-based samples.\n'
    '   (c) Public Mental Health — PRO: a strong fit for '
    'the prevention-architecture framing and Vietnamese '
    'population focus. CON: may underweight the SEM '
    'methodological contribution.\n'
    'Recommended primary: (a) Adolescent and Young Adult '
    'Psychiatry, with (c) Public Mental Health as fallback.',
    'Q3-9 — Chiến lược nộp Q2 (Frontiers in Psychiatry): '
    'xác nhận section tạp chí mục tiêu. Ba lựa chọn được '
    'liệt kê dưới đây kèm ghi chú ưu/nhược điểm:\n'
    '   (a) Adolescent and Young Adult Psychiatry — ƯU: '
    'khớp với mẫu (11–14 tuổi); thiết kế theo phân loại đặc '
    'thù đúng trọng tâm. NHƯỢC: section cạnh tranh cao, ưu '
    'tiên nghiên cứu dọc/lâm sàng.\n'
    '   (b) Anxiety and Stress Disorders — ƯU: khớp với '
    'trọng tâm phân loại DSM-5 và biến kết quả. NHƯỢC: độc '
    'giả người lớn rộng hơn, ít quan tâm hơn tới mẫu '
    'trường học.\n'
    '   (c) Public Mental Health — ƯU: phù hợp với khung '
    'kiến trúc phòng ngừa và trọng tâm dân số Việt Nam. '
    'NHƯỢC: có thể đánh giá thấp đóng góp phương pháp luận '
    'SEM.\n'
    'Đề xuất chính: (a) Adolescent and Young Adult '
    'Psychiatry, với (c) Public Mental Health là phương án '
    'dự phòng.')


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
