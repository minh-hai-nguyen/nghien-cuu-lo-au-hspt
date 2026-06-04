# -*- coding: utf-8 -*-
"""Viết tóm tắt batch — QT21-25, QT35, VN15-20"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'

sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '03_Ban-dich'))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables, {len([x for x in d.paragraphs if x.text.strip()])} paras')

def tt(title_vn, title_en, info, pp_method, pp_tools, pp_justify, data_demo,
       process, validity, results_table_header, results_table_rows, results_table_widths,
       compare, remarks, conclusion, critique, future, rating,
       extra_tables=None):
    """Tạo tóm tắt theo 13 bước CTH v5"""
    doc = create_doc()
    add_heading(doc, title_vn, 1)

    # 1. ĐỊNH DANH
    add_p(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
    add_p(doc, info)

    # 2. TỔNG QUAN PP
    add_p(doc, 'Phương pháp nghiên cứu', bold=True)
    add_p(doc, pp_method)

    # 3. ĐỊNH NGHĨA
    if pp_tools:
        add_p(doc, pp_tools)

    # 4. BIỆN MINH
    if pp_justify:
        add_p(doc, pp_justify)

    # 5-6. DỮ LIỆU + QUY TRÌNH
    if data_demo:
        add_p(doc, data_demo)
    if process:
        add_p(doc, process)

    # 7-8. CHỨNG THỰC
    if validity:
        add_p(doc, validity)

    # 9. ĐỊNH LƯỢNG
    add_p(doc, 'Kết quả nghiên cứu định lượng', bold=True)
    add_table(doc, results_table_header, results_table_rows, widths=results_table_widths)

    # Extra tables
    if extra_tables:
        for et in extra_tables:
            add_p(doc, '')
            add_table(doc, et['header'], et['rows'], widths=et.get('widths'))

    # 10. ĐỐI CHIẾU
    if compare:
        add_p(doc, 'Đối chiếu liên bài', bold=True)
        add_p(doc, compare)

    # 11-12. NHẬN XÉT
    add_p(doc, 'Nhận xét, phát hiện qua kết quả nghiên cứu', bold=True)
    for rm in remarks:
        add_p(doc, rm)

    # 13. KẾT LUẬN
    add_p(doc, 'Kết luận', bold=True)
    add_p(doc, conclusion)

    # PHẢN BIỆN
    add_red_heading(doc, 'Phản biện')
    add_red(doc, critique)

    # HƯỚNG NC
    add_red_heading(doc, 'Hướng nghiên cứu tiếp theo')
    add_red(doc, future)

    # ĐÁNH GIÁ
    add_p(doc, f'Đánh giá: {rating}', bold=True)

    return doc

# =====================================================================
# QT21 — Norway 2025
# =====================================================================
print('QT21...')
doc = tt(
    'Tóm tắt bài QT-21',
    'Possible Explanations for the Upward Trend in Mental Distress among Adolescents in Norway from 2011 to 2024',
    'Công trình « Giải thích khả thi cho xu hướng tăng căng thẳng tâm thần ở thanh thiếu niên Na Uy từ 2011 đến 2024 » (Possible Explanations for the Upward Trend in Mental Distress among Adolescents in Norway from 2011 to 2024), do Geir Scott Brunborg, Sondre Aasen Nilsen, Jens Christoffer Skogen và Lasse Bang (2025), phân tích 979.043 VTN Na Uy từ 1.417 khảo sát cấp thành phố (2011–2024). Social Science & Medicine, 384, 118528. DOI: 10.1016/j.socscimed.2025.118528.',
    'Công trình sử dụng phân tích cắt ngang lặp lại (repeated cross-sectional) với phương pháp decomposition — giữ cố định (holding constant) từng yếu tố để kiểm tra đóng góp vào xu hướng. Nói cách khác, đây là phương pháp "thí nghiệm tư duy dịch tễ học" — nếu loại bỏ tác động của một yếu tố, xu hướng tăng có biến mất không?',
    'Căng thẳng tâm thần đo bằng 6 mục rút gọn từ Thang Triệu chứng Hopkins (HSCL — Hopkins Symptom Checklist). Tám yếu tố giải thích: khó khăn tài chính, tối ở nhà, thiếu vận động, bắt nạt, bất mãn cha mẹ, mạng xã hội, cần sa, bất mãn trường học.',
    'Tổng quan tài liệu cho thấy xu hướng tăng căng thẳng tâm thần VTN đã được ghi nhận ở nhiều nước nhưng ít NC giải thích nguyên nhân. Khung Keyes & Platt (2024) đề xuất 3 tiêu chí: (1) liên quan cá nhân, (2) tỷ lệ tăng, (3) độc lực tăng.',
    'N = 979.043 VTN, chủ yếu lớp 8–10 (13–16 tuổi), từ 1.417 khảo sát cấp thành phố, 2011–2024 (trừ 2020).',
    'Hồi quy tuyến tính, decomposition (giữ cố định), kiểm tra tương tác yếu tố × thời gian. Stata 18.',
    None,
    ['Yếu tố', 'Tỷ lệ tăng?', 'Giữ cố định → xu hướng', 'Kết luận'],
    [['Bất mãn trường học', 'Có', 'LOẠI BỎ xu hướng (b→−0,004)', 'GIẢI THÍCH CHÍNH'],
     ['Mạng xã hội', 'Có', 'Giảm đáng kể (0,016→0,006)', 'Giải thích MỘT PHẦN'],
     ['Tối ở nhà', 'Có', 'Giảm nhẹ', 'Đóng góp nhỏ'],
     ['Cần sa', 'Có (nhỏ)', 'Giảm nhẹ (chỉ nữ)', 'Nhỏ, chỉ nữ'],
     ['Khó khăn tài chính', 'GIẢM', 'Không thay đổi', 'Không giải thích'],
     ['Hoạt động thể chất', 'TĂNG', 'Không thay đổi', 'Không giải thích'],
     ['Bắt nạt', 'Ổn định', 'Không thay đổi', 'Không giải thích'],
     ['Bất mãn cha mẹ', 'Ổn định', 'Không thay đổi', 'Không giải thích']],
    [3.5, 2.0, 4.0, 3.5],
    'Xu hướng 13 năm phù hợp với Norway 2025 (bài này), Korea 2024 (đảo chiều sau COVID), JAACAP 2024 (lo âu gấp đôi Mỹ). Bất mãn trường học phù hợp Cosma et al. (2020) ở 36 quốc gia.',
    ['*Bất mãn trường học — yếu tố giải thích chính.* Giữ cố định yếu tố này loại bỏ TOÀN BỘ xu hướng tăng ở nam. Cơ chế: áp lực học tập, kỳ vọng thành tích, thay đổi thực hành giáo dục.',
     '*Mạng xã hội — đóng góp một phần.* Phù hợp với Nature 2025 (Fassi et al.: VTN SKTT dùng MXH nhiều hơn, g = 0,46). Tuy nhiên, bằng chứng vẫn lẫn lộn — Li 2025 (BJCP) tìm thấy tác động dọc YẾU.',
     '*4 yếu tố KHÔNG giải thích xu hướng:* khó khăn tài chính (giảm), thể chất (tăng), bắt nạt (ổn định), bất mãn cha mẹ — rất rõ ràng.'],
    'Dữ liệu 979.043 VTN Na Uy qua 13 năm, cho thấy bất mãn trường học tăng và sử dụng mạng xã hội tăng là hai yếu tố giải thích chính cho xu hướng tăng căng thẳng tâm thần, gợi ý rằng can thiệp cần nhắm vào cải thiện môi trường trường học và quản lý thời gian mạng xã hội.',
    'Social Science & Medicine Q1 IF = 5,4. Mẫu cực lớn (979K). Decomposition nâng cao. Tuy nhiên, chỉ Na Uy — khác biệt văn hóa lớn với VN. Đo distress chung (HSCL), không tách lo âu riêng. Giả định nhân quả một chiều — hạn chế lớn nhất.',
    'Áp dụng decomposition cho dữ liệu VN (Hoàng Trung Học có 2 thời điểm). So sánh vai trò bất mãn trường học ở VN vs Na Uy. Can thiệp giảm bất mãn trường học + kiểm soát MXH tại VN.',
    '⭐⭐⭐⭐ Cao. Q1 IF = 5,4, n = 979K, 13 năm, decomposition.'
)
save(doc, 'QT21_Norway_2025.docx')

# =====================================================================
# QT22 — ScreenTime 2025 BJCP
# =====================================================================
print('QT22...')
doc = tt(
    'Tóm tắt bài QT-22',
    'Cross-sectional and Longitudinal Associations of Screen Time with Adolescent Depression and Anxiety',
    'Công trình « Mối liên quan cắt ngang và dọc giữa thời gian sử dụng màn hình với trầm cảm và lo âu ở thanh thiếu niên » (Cross-sectional and Longitudinal Associations of Screen Time with Adolescent Depression and Anxiety), do Sophie H. Li và cộng sự (2025), thuộc Black Dog Institute, Đại học New South Wales, Úc, phân tích 4.058 VTN Úc (tuổi TB = 13,9) từ 134 trường trung học, thuộc Nghiên cứu Future Proofing — thuần tập tiến cứu 5 năm. British Journal of Clinical Psychology, 64, 873–887. DOI: 10.1111/bjc.12547.',
    'Công trình sử dụng thiết kế DỌC (longitudinal) — theo dõi cùng nhóm VTN qua 12 tháng — kết hợp phân tích cắt ngang. Nói cách khác, đây là một trong số ít NC có thể kiểm tra chiều hướng nhân quả: screen time T1 có DỰ BÁO trầm cảm/lo âu T2 không?',
    'Trầm cảm: PHQ-A (Patient Health Questionnaire for Adolescents; α = 0,88), 0–27 điểm. Lo âu: CAS-8 (Children\'s Anxiety Scale Short-Form; α = 0,90), 0–24 điểm. Screen time: 1 câu tự báo cáo (0–1h đến 5+h). MXH tiêu cực: 7 mục (Smith et al., 2013), 7–49 điểm.',
    'Tổng quan tài liệu cho thấy đa số NC là cắt ngang — NC này lấp đầy khoảng trống bằng thiết kế dọc.',
    'Mẫu ban đầu 6.388 → 4.058 (loại thiếu dữ liệu). 134 trường trung học Úc. Nữ 52,3%, nam 44,3%, đa dạng giới 2,2%.',
    'Mô hình hỗn hợp tuyến tính (linear mixed models), trường = random effect. 3 mô hình cho mỗi biến.',
    None,
    ['Phân tích', 'Biến', 'Screen time b', 'p', 'Lâm sàng (5+h vs 0–1h)'],
    [['Cắt ngang T1', 'Trầm cảm PHQ-A', '1,25/giờ', '<0,001', '+6,25 điểm — ĐÁNG KỂ'],
     ['Cắt ngang T1', 'Lo âu CAS-8', '0,72/giờ', '<0,001', '+3,60 điểm — ĐÁNG KỂ'],
     ['DỌC T1→T2', 'Trầm cảm PHQ-A', '0,15/giờ', '0,007', '+0,75 điểm — RẤT NHỎ'],
     ['DỌC T1→T2', 'Lo âu CAS-8', '—', '0,443', 'KHÔNG có ý nghĩa']],
    [3.0, 3.0, 2.5, 1.5, 4.0],
    'Thiết kế dọc bổ sung cho Chen 2023 (game OR = 5,00 — cắt ngang), Hoàng Trung Học 2025 (điện tử Beta = 0,176 — cắt ngang), Norway 2025 (MXH giải thích xu hướng). Tuy nhiên, tác động dọc YẾU — gợi ý quan hệ HAI CHIỀU.',
    ['*Phát hiện then chốt và NGƯỢC LẠI giả định phổ biến:* Cắt ngang đáng kể nhưng dọc YẾU/KHÔNG đáng kể. Screen time T1 KHÔNG dự báo lo âu T2 (p = 0,443).',
     '*Gợi ý quan hệ HAI CHIỀU:* VTN có triệu chứng có thể TĂNG screen time (rút lui xã hội, phân tâm) — không chỉ ngược lại.',
     '*MXH tiêu cực KHÔNG ảnh hưởng đáng kể* — có thể do thang đo Facebook 2013 lỗi thời.',
     '*Giới tính:* Nữ liên quan mạnh hơn nam với screen time → trầm cảm (cắt ngang), nhưng không khác biệt dọc.'],
    'Dữ liệu dọc 4.058 VTN Úc, cho thấy screen time có liên quan cắt ngang đáng kể nhưng chỉ liên quan dọc yếu với trầm cảm và KHÔNG liên quan dọc với lo âu, gợi ý rằng mối quan hệ phức tạp hơn giả định — có thể hai chiều, và cần thận trọng khi kết luận screen time GÂY RA giảm SKTT.',
    'BJCP Q1. Thiết kế DỌC — hiếm. Mẫu lớn (4.058), 134 trường, Future Proofing Study. PHQ-A α=0,88, CAS-8 α=0,90. Tuy nhiên, screen time chỉ 1 câu tự báo cáo. Chỉ Úc. 12 tháng có thể quá ngắn. Thang MXH lỗi thời (Facebook 2013).',
    'NC dọc tương tự tại VN. Phân biệt LOẠI screen time (MXH, game, học tập). RCT giảm screen time tại trường VN — kết hợp JAMA 2024.',
    '⭐⭐⭐⭐ Cao. BJCP Q1, thiết kế dọc, n = 4.058.'
)
save(doc, 'QT22_ScreenTime_2025.docx')

# =====================================================================
# QT23 — JAACAP US 2024
# =====================================================================
print('QT23...')
doc = tt(
    'Tóm tắt bài QT-23',
    'Trends in Mental Disorders in Children and Adolescents Receiving Treatment in the State Mental Health System',
    'Công trình « Xu hướng rối loạn tâm thần ở trẻ em và VTN nhận điều trị trong hệ thống SKTT công lập Hoa Kỳ » (Trends in Mental Disorders in Children and Adolescents Receiving Treatment in the State Mental Health System), do Ramin Mojtabai và Mark Olfson (2024), Tulane University và Columbia University, phân tích 13.684.154 hồ sơ trẻ em/VTN (≤17 tuổi) trong hệ thống SKTT công Mỹ (MH-CLD), 2013–2021. JAACAP, Vol. 64(8). DOI: 10.1016/j.jaac.2024.08.008.',
    'Phân tích xu hướng tỷ lệ các rối loạn tâm thần được chẩn đoán trong hệ thống y tế công lập Mỹ, sử dụng dữ liệu hành chính quốc gia 8 năm. Nói cách khác, đây là phân tích xu hướng dựa trên chẩn đoán LÂM SÀNG (không phải sàng lọc) từ hệ thống dịch vụ thực tế.',
    'Dữ liệu MH-CLD (Mental Health Client-Level Data) — SAMHSA. 8 nhóm rối loạn: lo âu, trầm cảm, lưỡng cực, sang chấn/stress, ADHD, hành vi, chống đối, phát triển lan tỏa.',
    'Dữ liệu chẩn đoán lâm sàng (vị trí 1-3 trong hồ sơ). Khác với sàng lọc — phản ánh tỷ lệ ĐƯỢC CHẨN ĐOÁN thực tế.',
    'N = 13.684.154 hồ sơ, 2013–2021. Mỹ 1,6–1,9 triệu/năm. 50 tiểu bang + DC.',
    'Hồi quy logistic (AOR), điều chỉnh tuổi, giới, chủng tộc. Clustering tiểu bang. Linear splines kiểm tra COVID.',
    None,
    ['Rối loạn', '2013 (%)', '2021 (%)', 'AOR (KTC 95%)', 'Xu hướng'],
    [['Rối loạn lo âu', '9,6', '19,2', '2,17 (1,85–2,55)***', '↑ TĂNG GẤP ĐÔI'],
     ['Rối loạn trầm cảm', '13,4', '17,0', '1,20 (1,03–1,41)*', '↑ Tăng'],
     ['Sang chấn/stress', '22,7', '27,4', '1,31 (1,09–1,57)**', '↑ Tăng'],
     ['ADHD', '29,8', '25,3', '0,90 (0,77–1,06)', '— Ổn định'],
     ['Rối loạn lưỡng cực', '10,0', '1,3', '0,07 (0,06–0,09)***', '↓ GIẢM 8 LẦN'],
     ['Rối loạn hành vi', '9,7', '4,4', '0,42 (0,33–0,55)***', '↓ Giảm mạnh']],
    [3.5, 1.5, 1.5, 3.5, 3.0],
    'Lo âu tăng gấp đôi (AOR = 2,17) — xu hướng nhất quán toàn cầu: Norway 2025 (13 năm tăng), GBD 2025 (AAPC 0,84%), Korea 2024 (tăng sau COVID). So với V-NAMHS 2022 (chỉ 2,3% chẩn đoán) — khoảng cách khổng lồ.',
    ['*Lo âu — rối loạn tăng nhanh nhất.* AOR = 2,17 vượt trội trầm cảm (1,20) và sang chấn (1,31). Lo âu đang trở thành cuộc khủng hoảng SKTT hàng đầu ở trẻ em/VTN phương Tây.',
     '*Lo âu tăng mạnh nhất ở 15–17 tuổi* (AOR = 2,93 — gần gấp BA).',
     '*Lưỡng cực giảm 8 lần* — phản ánh thay đổi thực hành chẩn đoán (DSM-5, DMDD).',
     '*Dữ liệu chẩn đoán lâm sàng* — khác với sàng lọc. So với Hoa 2024 (40,6% GAD-7) và V-NAMHS (2,3% DISC-5): chênh lệch phản ánh khác biệt phương pháp đo.'],
    'Dữ liệu từ 13,7 triệu hồ sơ lâm sàng Mỹ, cho thấy lo âu tăng gấp đôi (AOR = 2,17) trong 8 năm — xu hướng mạnh nhất, đặc biệt ở 15–17 tuổi, gợi ý rằng lo âu đang trở thành ưu tiên SKTT hàng đầu toàn cầu. Tại VN, nơi chỉ 8,4% VTN tiếp cận dịch vụ, tỷ lệ thực có thể cao hơn nhiều so với báo cáo.',
    'JAACAP Q1 IF ≈ 11. N = 13,7 triệu hồ sơ lâm sàng. Xu hướng 8 năm. Phân tầng tuổi/giới/chủng tộc. Tuy nhiên: chỉ trẻ ĐANG điều trị (thiên lệch chọn), không tách tăng thực vs tăng phát hiện, chỉ Mỹ.',
    'Phân tích xu hướng tương tự từ BV Nhi TW VN. So sánh Mỹ (82,6% tiếp cận) vs VN (8,4%). Đánh giá nguyên nhân xu hướng tăng.',
    '⭐⭐⭐⭐⭐ Rất cao. JAACAP Q1 IF = 11, n = 13,7 triệu, chẩn đoán lâm sàng.'
)
save(doc, 'QT23_JAACAP_US_2024.docx')

# =====================================================================
# QT24 — WHO Europe 2025
# =====================================================================
print('QT24...')
doc = tt(
    'Tóm tắt bài QT-24',
    'Mental Health of Children and Young People in the WHO Europe Region',
    'Công trình « Sức khỏe tâm thần của trẻ em và thanh niên trong khu vực WHO châu Âu » (Mental Health of Children and Young People in the WHO Europe Region), do Anna Tarasenko và cộng sự (2025), xuất bản trên The Lancet Regional Health — Europe (Q1, IF ≈ 15,0). Bài Series "Transforming Mental Health in Europe". 13 trang.',
    'Tổng quan chính sách (policy review) tổng hợp dữ liệu từ nhiều nguồn trong khu vực WHO châu Âu (53 quốc gia). Nói cách khác, đây KHÔNG phải nghiên cứu gốc mà là bài phân tích tổng hợp phục vụ hoạch định chính sách.',
    None,
    'Tổng quan tài liệu cho thấy 9 triệu VTN châu Âu (10–19 tuổi) sống với rối loạn SKTT. Lo âu + trầm cảm chiếm >50% tổng ca.',
    None,
    None,
    None,
    ['Lĩnh vực', 'Dữ liệu', 'Khuyến nghị', 'Ý nghĩa cho VN'],
    [['Tỷ lệ', '9 triệu VTN có RLSKTT', 'Ưu tiên quốc gia', 'VN: chưa có dữ liệu tổng'],
     ['Lo âu + trầm cảm', '>50% tổng ca rối loạn', 'Can thiệp tập trung', 'Phù hợp — lo âu hàng đầu'],
     ['Trường học', 'SKTT vào giáo dục', 'MindMatters, SEAL', 'Wen 2020: OR = 0,562 bảo vệ'],
     ['MXH', 'Yếu tố nguy cơ', 'Kiểm soát nội dung', 'Norway 2025: giải thích xu hướng'],
     ['Dịch vụ', '16/27 nước thiếu', 'Phát triển cộng đồng', 'VN: 8,4% tiếp cận'],
     ['Nghịch cảnh', 'ACEs → SKTT', 'Sàng lọc + can thiệp', 'Ngô Anh Vinh 2024: ACEs ở DTTS']],
    [2.5, 3.5, 3.0, 4.5],
    'Con số 9 triệu VTN châu Âu cho thấy quy mô vấn đề. So sánh: ASEAN 80,4 triệu ca rối loạn (GBD 2025). Tại VN, chỉ 8,4% tiếp cận dịch vụ vs hệ thống phát triển của châu Âu.',
    ['*Khuyến nghị tích hợp SKTT vào giáo dục* — phù hợp Wen 2020 (hỗ trợ trường OR = 0,562) và Zhameden 2025 (can thiệp trường LMIC).',
     '*Lo âu + trầm cảm > 50% ca rối loạn* — xác nhận lo âu là vấn đề hàng đầu, phù hợp JAACAP 2024 và tất cả NC VN.',
     '*Chỉ 5/53 nước có tổ chức đa ngành cho SKTT* — ngay cả châu Âu cũng thiếu. VN cần học từ mô hình Bắc Âu (bottom-up).'],
    'Dữ liệu tổng hợp WHO châu Âu, cho thấy 9 triệu VTN đang sống với rối loạn SKTT và lo âu chiếm > 50% tổng ca, gợi ý rằng tích hợp SKTT vào hệ thống giáo dục và y tế là chiến lược thiết yếu — bài học cho VN trong bối cảnh xây dựng chương trình SKTT học đường.',
    'Lancet Regional Health Europe Q1 IF = 15. Phạm vi 53 quốc gia. Bài Series chính sách. Tuy nhiên: KHÔNG phải NC gốc — thiếu số liệu chi tiết từng nước. Bối cảnh châu Âu (dịch vụ phát triển) khác VN.',
    'Tổng quan tương tự cho ASEAN/VN. So sánh mô hình châu Âu vs VN. Đánh giá chi phí–hiệu quả can thiệp SKTT trường cho VN.',
    '⭐⭐⭐⭐⭐ Rất cao. Lancet Q1 IF = 15, 53 quốc gia, hướng chính sách.'
)
save(doc, 'QT24_WHO_Europe_2025.docx')

# =====================================================================
# QT25 — EpiPsychSci 2025
# =====================================================================
print('QT25...')
doc = tt(
    'Tóm tắt bài QT-25',
    'The Complete Mental Health of Australia\'s Adolescents and Emerging Adults',
    'Công trình « Sức khỏe tâm thần toàn diện của VTN và người trưởng thành mới nổi tại Úc » (The Complete Mental Health of Australia\'s Adolescents and Emerging Adults), do Dimity Crisp, Debra Rickwood, Richard Burns và Emily Bariola (2025), Đại học Canberra và Đại học Quốc gia Úc, 3 khảo sát quốc gia: 2018 (n=3.721), 2020 (n=974), 2022 (n=961), VTN 12–25 tuổi. Epidemiology and Psychiatric Sciences, 34, e16. DOI: 10.1017/S2045796025000083.',
    'Sử dụng khung Sức khỏe Tâm thần Toàn diện (CMH — Complete Mental Health) của Keyes — đo CẢ bệnh tật (K10) VÀ hạnh phúc (MHC-SF). Nói cách khác, thay vì chỉ hỏi "bạn có bệnh không?", đo thêm "bạn có hạnh phúc không?", tạo ra 6 trạng thái SKTT.',
    'K10 (Kessler-10): 10 mục, đo căng thẳng tâm lý. MHC-SF (Mental Health Continuum Short Form): 14 mục, đo hạnh phúc. 6 trạng thái CMH: Flourishing, Middling, Languishing, Struggling, Stumbling, Floundering.',
    'Tổng quan cho thấy NC SKTT chủ yếu tập trung vào bệnh tật — thiếu đánh giá SKTT tích cực.',
    'Thiết kế cắt ngang lặp lại. 3 khảo sát quốc gia: 2018 (n=3.721), 2020 (n=974), 2022 (n=961). Phỏng vấn điện thoại CATI. Mẫu quota.',
    'ANOVA, hồi quy đa thức. Phân tích theo năm, tuổi, giới.',
    None,
    ['Trạng thái', '2018 (%)', '2022 (%)', 'Xu hướng', 'Ý nghĩa'],
    [['Flourishing', '53,0', '44,4', '↓ GIẢM (p<0,001)', 'Ít hạnh phúc + không bệnh'],
     ['Stumbling', '19,3', '27,3', '↑ TĂNG MẠNH (p<0,001)', 'TB hạnh phúc + có bệnh'],
     ['Struggling', '9,1', '10,2', '— Ổn định', 'Hạnh phúc + có bệnh'],
     ['K10 cao/rất cao', '30,7', '40,7', '↑ TĂNG', 'Căng thẳng tâm lý']],
    [3.0, 2.0, 2.0, 3.5, 3.5],
    'Xu hướng flourishing giảm phù hợp: Norway 2025 (distress tăng 13 năm), JAACAP 2024 (lo âu gấp đôi), Ireland 2024 (trầm cảm/lo âu tăng 2012–2019). Khoảng cách giới mở rộng phù hợp xu hướng toàn cầu.',
    ['*Flourishing giảm 53%→44,4%* — SKTT giới trẻ xấu đi CẢ VỀ tăng bệnh VÀ giảm hạnh phúc.',
     '*Khoảng cách giới MỞ RỘNG:* Nam flourishing 52,2% vs nữ 36,3% (2022) — chênh 15,9 điểm, tăng từ 10,4 điểm (2018).',
     '*Flourishing giảm theo tuổi:* 70,7% ở 12–14 → 43,3% ở 18–21 — giai đoạn chuyển tiếp quan trọng.',
     '*Mô hình CMH phát hiện "struggling":* 9–10% có bệnh nhưng VẪN hạnh phúc — hạnh phúc có thể là yếu tố bảo vệ.'],
    'Dữ liệu 3 khảo sát quốc gia Úc, cho thấy SKTT toàn diện giới trẻ xấu đi đáng kể (flourishing giảm, stumbling tăng, K10 cao tăng lên 40,7%), gợi ý rằng cần tiếp cận CMH toàn diện — đo cả bệnh tật và hạnh phúc — thay vì chỉ tập trung vào pathology.',
    'EpiPsychSci Q1 IF ≈ 7. CMH toàn diện — đo cả bệnh VÀ hạnh phúc. 3 khảo sát. Tuy nhiên: chỉ Úc. Mẫu 2020-2022 nhỏ. K10 không tách lo âu riêng. CATI — có thể thiên lệch.',
    'Áp dụng khung CMH cho VTN VN. Đo hạnh phúc song song bệnh tật. So sánh flourishing Úc vs VN.',
    '⭐⭐⭐⭐ Cao. Q1 IF = 7, CMH toàn diện, 3 khảo sát quốc gia.',
    extra_tables=[{'header': ['Giới', 'Flourishing 2018', 'Flourishing 2022', 'Chênh lệch'], 'rows': [['Nam','58,3%','52,2%','↓ 6,1 điểm'],['Nữ','47,9%','36,3%','↓ 11,6 điểm'],['Khoảng cách','10,4 điểm','15,9 điểm','MỞ RỘNG']], 'widths': [3.0, 3.0, 3.0, 3.0]}]
)
save(doc, 'QT25_EpiPsychSci_2025.docx')

# =====================================================================
# QT35 — Social Anxiety 7 Countries
# =====================================================================
print('QT35...')
doc = tt(
    'Tóm tắt bài QT-35',
    'Social Anxiety in Young People: A Prevalence Study in Seven Countries',
    'Công trình « Lo âu xã hội ở thanh niên: Nghiên cứu tỷ lệ tại 7 quốc gia » (Social Anxiety in Young People: A Prevalence Study in Seven Countries), do Philip Jefferies và Michael Ungar (2020), Dalhousie University, Canada, khảo sát 6.825 thanh niên 16–29 tuổi (TB = 22,84; SD = 3,97) từ Brazil, Trung Quốc, Indonesia, Nga, Thái Lan, Mỹ, và VIỆT NAM. PLOS ONE, 15(9), e0239133. DOI: 10.1371/journal.pone.0239133.',
    'Sử dụng Thang Lo âu Tương tác Xã hội 17 mục (SIAS-17 — Social Interaction Anxiety Scale; Mattick & Clarke, 1998) đo nỗi sợ và lo âu liên quan tương tác xã hội. Nói cách khác, đây đo lo âu XÃ HỘI cụ thể (không phải lo âu tổng quát GAD) — sợ bị đánh giá, sợ nói chuyện, sợ tương tác.',
    'SIAS-17: 17 mục thuận chiều, ngưỡng SAD (Social Anxiety Disorder) ≥ 29 (tương đương SIAS-20 ≥ 34). Phân biệt lâm sàng/không lâm sàng, phân biệt SAD vs GAD. Đã kiểm chứng đa văn hóa.',
    'Tổng quan cho thấy tỷ lệ SAD suốt đời lên đến 12% ở Mỹ, nhưng các NC trước chủ yếu đo ở người trưởng thành, ít tập trung thanh niên và đa quốc gia.',
    '6.825 người, mẫu quota từ panel đại diện quốc gia (Dynata, OMI, GMO Research). Khảo sát trực tuyến 11/2019.',
    'ANOVA, chi-square, hồi quy. So sánh theo giới, tuổi, quốc gia, việc làm, khu vực, học vấn. Đánh giá tự nhận thức vs ngưỡng SIAS.',
    None,
    ['Quốc gia', 'Điểm TB (SD)', 'SAD ≥ 29 (%)', 'So sánh'],
    [['Brazil', '26,18 (15,23)', '42,4%', ''],
     ['Trung Quốc', '22,30 (13,52)', '32,1%', ''],
     ['Indonesia', '18,94 (13,21)', '22,9%', 'Thấp nhất'],
     ['Nga', '20,78 (12,79)', '27,0%', ''],
     ['Thái Lan', '25,57 (13,92)', '41,4%', ''],
     ['Hoa Kỳ', '30,35 (15,44)', '57,6%', 'Cao nhất'],
     ['VIỆT NAM', '22,68 (11,77)', '30,7%', 'SD thấp nhất'],
     ['TỔNG', '23,82 (14,18)', '36,2%', 'F = 74,85, p < 0,001']],
    [3.0, 3.5, 2.5, 4.0],
    'VN SAD = 30,7% — cao hơn nhiều so với V-NAMHS 2022 (2,3% chẩn đoán DISC) — do đo khác (sàng lọc vs chẩn đoán). So với bài khác: Hoa 2024 (40,6% GAD-7 lo âu tổng quát), Ngô Anh Vinh 2024 (lo âu DASS-21 ở DTTS). Đây là NC DUY NHẤT trong Đề tài đo lo âu XÃ HỘI riêng ở VN.',
    ['*Tỷ lệ SAD = 36% toàn cầu — cao hơn nhiều so với 12% báo cáo trước.* Do: (1) tập trung thanh niên 16–29, (2) thay đổi xã hội (MXH), (3) tăng nhận thức.',
     '*KHÔNG khác biệt giới tính* — trái kỳ vọng. Lo âu xã hội có thể ảnh hưởng nam=nữ (khác GAD thường nữ>nam).',
     '*18–24 tuổi cao nhất (40,3%)* — giai đoạn đại học/đầu sự nghiệp = đỉnh lo âu xã hội.',
     '*18% "false negatives"* — có SAD nhưng KHÔNG nhận ra. Tại VN, tỷ lệ này có thể cao hơn do nhận thức SKTT còn thấp.',
     '*VN: nói với người có thẩm quyền là lo âu xếp hạng 3* — phù hợp văn hóa tôn trọng thứ bậc.'],
    'Dữ liệu 6.825 thanh niên từ 7 quốc gia, cho thấy lo âu xã hội rất phổ biến (36% đạt ngưỡng SAD), không khác biệt giới tính, cao nhất 18–24 tuổi. VIỆT NAM: 30,7% — cứ 3 thanh niên VN có 1 người có khả năng mắc SAD. 18% là "lo âu ẩn". Gợi ý rằng cần sàng lọc lo âu XÃ HỘI riêng — không chỉ lo âu tổng quát.',
    'PLOS ONE Q1. 7 quốc gia đa dạng, BÀO GỒM VN. Mẫu lớn (6.825). SIAS đa văn hóa. Tuy nhiên: mẫu market research (Unilever), không phải khảo sát SKTT. SIAS chỉ đo tương tác (không đo biểu diễn). Sàng lọc, không chẩn đoán. Tuổi 16–29 rộng hơn VTN.',
    'NC lo âu xã hội chuyên biệt ở VTN THCS/THPT VN (10–17 tuổi). So sánh VN vs Indonesia, Thái Lan (cùng ĐNA nhưng rất khác). Đánh giá "false negatives" ở VN — nhận thức SKTT. MXH + lo âu XH: NC riêng.',
    '⭐⭐⭐⭐ Cao. PLOS ONE Q1, 7 quốc gia bao gồm VN, n = 6.825.',
    extra_tables=[{'header': ['Phân nhóm', 'SAD (%)', 'p', 'Ghi chú'], 'rows': [['Nam','35,6%','n.s.','KHÔNG khác biệt giới'],['Nữ','36,5%','',''],['16–17 tuổi','30,8%','<0,001',''],['18–24 tuổi','40,3%','','Cao nhất'],['25–29 tuổi','32,8%','',''],['Thất nghiệp','41,7%','<0,001','Cao nhất'],['Chưa xong cấp 3','52,0%','<0,001','Rất cao'],['Ngoại ô','42,4%','<0,001','Cao nhất (không phải nông thôn)']], 'widths': [3.5, 2.0, 2.0, 4.0]}]
)
save(doc, 'QT35_SocialAnxiety_7Countries.docx')

print('\n=== DONE 5 TOM TAT ===')
