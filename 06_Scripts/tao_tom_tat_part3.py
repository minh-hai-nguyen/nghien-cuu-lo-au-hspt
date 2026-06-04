# -*- coding: utf-8 -*-
"""Tao tom tat bai 50-58 (9 bai cuoi cung)"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w * 567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
def save(doc, name):
    p = os.path.join(OUT_DIR, name); doc.save(p)
    d = Document(p); chars = sum(len(x.text) for x in d.paragraphs)
    print(f'  {name}: {chars} chars, {len(d.tables)} tables')

# ============================================================
# BÀI 50 — Cao 2025 Resilience MA
# ============================================================
print('[12/20] Bài 50: Cao 2025 Resilience MA')
doc = make_doc()
H(doc, 'Tóm tắt bài QT44 — SR + MA can thiệp resilience tại trường', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"School-based interventions for resilience in children and adolescents: a systematic review and meta-analysis of '
       'randomized controlled trials" của Cao, C. và cộng sự (2025), đăng trên Frontiers in Psychiatry vol. 16, article '
       '1594658, tr. 1–15. DOI 10.3389/fpsyt.2025.1594658. Q1, IF ≈ 4,7. Open Access. Tổng quan hệ thống và phân tích '
       'tổng hợp các RCT về can thiệp khả năng phục hồi (resilience) tại trường cho trẻ em + vị thành niên.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống các RCT về can thiệp resilience tại trường. Đánh giá chất lượng bằng Cochrane Risk of Bias. '
       'Random-effects meta-analysis. Resilience = khả năng thích ứng tốt trước nghịch cảnh, căng thẳng hoặc sang chấn. '
       'Khác CBT (giảm triệu chứng) — tăng yếu tố BẢO VỆ: tự trọng, kỹ năng ứng phó, kết nối xã hội, lạc quan.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Can thiệp resilience tại trường CÓ HIỆU QUẢ tăng resilience và giảm triệu chứng SKTT ở trẻ em/VTN. Tuy nhiên, '
       'kích thước hiệu ứng NHỎ–TRUNG BÌNH và heterogeneity CAO. Nghĩa là: hiệu quả tồn tại nhưng không lớn, và có sự '
       'khác biệt đáng kể giữa các nghiên cứu (do khác biệt thiết kế, mẫu, văn hoá).')
H(doc, 'Phản biện', level=2)
P(doc, 'Frontiers Q1 IF ≈ 4,7, Open Access PMC. SR + MA RCTs — bằng chứng tốt. Resilience là yếu tố BẢO VỆ — bổ sung '
       'cho CBT, không thay thế. Phù hợp với Trần Thảo Vi 2025 (lạc quan trung gian, β = −0,24) và Ireland MyWorld 2024 '
       '(resilience + tự trọng quan trọng tăng theo thời gian). Hạn chế: hiệu lực nhỏ–TB; heterogeneity cao; đa số RCT '
       'phương Tây.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Phát triển module resilience tích hợp vào can thiệp CBT trường VN. RCT VN cho VTN có yếu tố nguy cơ cao. Thiết '
       'kế đo cả triệu chứng (CBT) lẫn nguồn lực bảo vệ (resilience).')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SR+MA Q1, định hướng module bổ trợ.', bold=True)
save(doc, 'QT44_Cao_Resilience_MA_2025.docx')

# ============================================================
# BÀI 51 — Sasaki 2024 JMIR Pediatrics Japan iCBT
# ============================================================
print('[13/20] Bài 51: Sasaki 2024 Japan iCBT subthreshold SAD')
doc = make_doc()
H(doc, 'Tóm tắt bài QT45 — RCT iCBT đa trung tâm cho subthreshold SAD tại Nhật Bản', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for Subthreshold Social Anxiety Disorder '
       'in Adolescents and Young Adults: Multicenter Randomized Controlled Trial" của Sasaki, N. và cộng sự (2024), đăng '
       'trên JMIR Pediatrics and Parenting vol. 7, e55786, tr. 1–13. DOI 10.2196/55786. UMIN000049768. Đối tượng: 77 học '
       'sinh / sinh viên có triệu chứng SAD DƯỚI NGƯỠNG (subthreshold) tại 6 đại học và 1 trường THPT Nhật Bản, từ 12/2022 '
       'đến 10/2023.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'RCT đa trung tâm 2 nhánh: nhóm can thiệp (n = 38) nhận iCBT TỰ HỌC HOÀN TOÀN qua web, không có hỗ trợ người; nhóm '
       'đối chứng (n = 39) chờ đợi. Can thiệp 8 module, theo dõi đến tháng 3 sau khởi động. Đo lường: SIAS (lo âu xã hội '
       'chính), PHQ-9 (trầm cảm). Phân tích ANCOVA với trầm cảm là biến đồng.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Nhóm iCBT giảm có ý nghĩa thống kê triệu chứng lo âu xã hội so với nhóm đối chứng. Tỷ lệ đáp ứng (response rate) '
       '61 % với OR = 4,97 so với chờ đợi. Tỷ lệ phục hồi (recovery rate) 68 % với OR = 3,95.')
table(doc, ['Chỉ số', 'Nhóm can thiệp', 'OR vs chờ', 'p'],
      [['Đáp ứng', '61 %', '4,97', '< 0,01'],
       ['Phục hồi', '68 %', '3,95', '< 0,01']],
      widths=[5.0, 4.0, 3.0, 3.5])
H(doc, 'Phản biện', level=2)
P(doc, 'RCT đa trung tâm — bằng chứng cao. Đặc biệt quan trọng vì SAD DƯỚI NGƯỠNG là dân số rất lớn và quan trọng — can '
       'thiệp ở giai đoạn này có thể NGĂN CHẶN tiến triển thành rối loạn đầy đủ ("indicated prevention"). Mô hình iCBT '
       'TỰ HỌC HOÀN TOÀN có chi phí gần bằng 0 sau khi phát triển nội dung — phù hợp triển khai đại trà cho VN. Mâu thuẫn '
       'với Walder 2025 (DMHI không hướng dẫn g chỉ ~ 0,3) — có thể do nội dung được thiết kế tốt hoặc effect bị overestimate '
       'do mẫu nhỏ. Hạn chế: n = 77 nhỏ, 1 quốc gia.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Adapt iCBT Nhật cho VN — sàng lọc HS subthreshold SAD bằng SIAS-17, can thiệp web tiếng Việt. So sánh có/không '
       'hỗ trợ người trong bối cảnh VN. RCT lớn hơn (n > 200) để xác nhận hiệu quả.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. RCT đa trung tâm Q2, mô hình prevention sớm khả thi.', bold=True)
save(doc, 'QT45_Sasaki_Japan_iCBT_2024.docx')

# ============================================================
# BÀI 52 — Academic Stress Interventions SR (2024/2025)
# ============================================================
print('[14/20] Bài 52: Academic Stress Interventions SR 2025')
doc = make_doc()
H(doc, 'Tóm tắt bài QT46 — Tổng quan hệ thống can thiệp căng thẳng học tập ở HS trung học', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Academic Stress Interventions in High Schools: A Systematic Literature Review" — đăng trên Child Psychiatry and '
       'Human Development vol. 56, tr. 1836–1869 (Q1, IF ≈ 3,3) — Springer, 2025. DOI 10.1007/s10578-024-01667-5. Tổng '
       'quan tài liệu hệ thống các nghiên cứu can thiệp giảm căng thẳng học tập (academic stress) tại trường THPT trên '
       'toàn cầu. 34 trang — toàn diện nhất hiện có về chủ đề này.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống các nghiên cứu can thiệp riêng cho stress học tập ở HS THPT (CBT cho stress học tập, mindfulness, '
       'kỹ năng quản lý thời gian, thư giãn, yoga, ...) trên toàn cầu. Phân loại theo loại can thiệp, hiệu quả và bối cảnh.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Đây là tổng quan hệ thống DUY NHẤT về can thiệp stress học tập riêng ở HS trung học. Stress học tập = yếu tố MẠNH '
       'NHẤT ảnh hưởng SKTT VTN châu Á — đã được xác nhận bởi VN21 Trần Thảo Vi (β = 4,73 cho học thêm), Wen QT08 '
       '(OR = 11,58 cho áp lực HT), Norway QT21, UNICEF VN22 (47 % HS học thêm > 3h/tuần).')
table(doc, ['Loại can thiệp', 'Hiệu quả', 'Phù hợp VN'],
      [['CBT cho stress HT', 'CAO nhất', 'Cần thích ứng VH'],
       ['Mindfulness/thiền', 'TB — kết quả lẫn lộn', 'UK 8.376 HS thất bại'],
       ['Quản lý thời gian', 'Hứa hẹn', 'Rất phù hợp VN'],
       ['Thư giãn cơ', 'TB', 'Dễ triển khai'],
       ['Yoga/hoạt động thể chất', 'Hứa hẹn', 'Phù hợp PE'],
       ['Kỹ năng ứng phó', 'TB–Khá', 'Kết hợp CBT']],
      widths=[5.0, 4.5, 6.0])
H(doc, 'Phản biện', level=2)
P(doc, 'Child Psychiatry Q1 IF ≈ 3,3, Springer. 34 trang — rất toàn diện. Tổng quan DUY NHẤT về stress HT — rất phù hợp '
       'đề tài VN. CBT hiệu quả nhất — phù hợp QT29 BMC NMA. Quản lý thời gian rất phù hợp VN (HS học thêm quá nhiều). '
       'Hạn chế: đa số nghiên cứu phương Tây, thiếu LMIC/châu Á; không có meta-analysis định lượng.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'RCT can thiệp stress HT riêng tại VN — CBT + quản lý thời gian + giảm tải học thêm. So sánh trong bối cảnh học '
       'sinh VN (sức ép đa chiều: cha mẹ, thi đại học, học thêm). Đánh giá hiệu quả dài hạn 1–2 năm.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SLR DUY NHẤT, Q1 Springer, rất phù hợp đề tài.', bold=True)
save(doc, 'QT46_AcademicStress_SR_2025.docx')

# ============================================================
# BÀI 53 — Dinh 2021 VN School Factors
# ============================================================
print('[15/20] Bài 53: Dinh 2021 VN School Factors')
doc = make_doc()
H(doc, 'Tóm tắt bài VN27 — Yếu tố trường học liên quan lo âu HS THCS Việt Nam', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "School factors causing Vietnamese adolescents anxiety" của Dinh và cộng sự (2021), đăng trên ResearchGate '
       'năm 2021. Khảo sát các yếu tố trường học liên quan lo âu ở HS THCS Việt Nam.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Cắt ngang trên HS THCS tại Việt Nam. Sử dụng bảng hỏi về yếu tố trường (áp lực học tập, mối quan hệ với GV và bạn '
       'bè, môi trường lớp học, ...) và thang đo lo âu. Phân tích hồi quy đa biến.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Các yếu tố trường học có liên quan có ý nghĩa với lo âu HS THCS Việt Nam: áp lực học tập, kỳ vọng từ GV, mối '
       'quan hệ tiêu cực với bạn bè (bạo lực học đường, bắt nạt), khối lượng bài tập về nhà nhiều, ít thời gian nghỉ ngơi.')
H(doc, 'Phản biện', level=2)
P(doc, 'Bài đăng trên ResearchGate (preprint hoặc tạp chí trong nước), không phải tạp chí Q1. Cung cấp dữ liệu định '
       'lượng về yếu tố trường học cụ thể tại VN — bổ sung cho UNICEF VN22 (School Factors). Hạn chế: không phải peer-review '
       'cấp cao, có thể có sai lệch chọn mẫu, không có thông tin chi tiết về cỡ mẫu và công cụ trong bài tóm lược.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Lặp lại với mẫu lớn, công cụ chuẩn (DASS-Y hoặc GAD-7), đăng tạp chí Q1/Q2. Phân tích mediation: yếu tố trường → '
       'cơ chế tâm lý → lo âu.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐ Trung bình. Có giá trị bổ sung dữ liệu trường học VN.', bold=True)
save(doc, 'VN27_Dinh_SchoolFactors_2021.docx')

# ============================================================
# BÀI 54 — Dong 2025 PLOS DASS-21 Ya'an
# ============================================================
print('[16/20] Bài 54: Dong 2025 PLOS Ya An TQ DASS-21')
doc = make_doc()
H(doc, 'Tóm tắt bài QT47 — Tỷ lệ và yếu tố quyết định trầm cảm/lo âu/stress ở HS trung học Ya An TQ', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Prevalence and determinants of depression, anxiety, and stress among secondary school students" của Dong, T., '
       'Wang, Y. & Lin, Y. (2025), đăng trên PLOS ONE vol. 20(9), e0328785, tr. 1–10. DOI 10.1371/journal.pone.0328785. '
       'Open Access. Mẫu: 2.716 học sinh trung học (1.230 nam + 1.486 nữ) tại Ya An, Tứ Xuyên, Trung Quốc, khảo sát điện '
       'tử tháng 3–6/2022. Sử dụng thang DASS-21.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Cắt ngang qua khảo sát điện tử (QuestionStar) tại 4 trường trung học Ya An. Sampling cụm, mẫu mục tiêu n = 1.844 '
       '(DE = 4, attrition 20 %), thu được 2.716 phiếu hợp lệ (93,66 %). DASS-21 đo cả 3 trục: trầm cảm, lo âu, stress. '
       'Bảng hỏi nhân khẩu học. Hồi quy logistic Model 1 đơn biến → Model 2 đa biến (loại p > 0,1 hoặc VIF > 10). Phần '
       'mềm R, p < 0,05 hai phía.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ: trầm cảm 24,4 %, LO ÂU 41,4 % (cao hơn rõ rệt), stress 15,6 %. Lo âu nặng/rất nặng gần 9 %. Hồi quy đa '
       'biến (Model 2) phát hiện CƠ CHẾ BẢO VỆ rõ ràng:')
table(doc, ['Yếu tố bảo vệ', 'OR (Trầm cảm)', 'OR (Lo âu)', 'Ý nghĩa'],
      [['Tâm sự với gia đình (vs không)', '0,22 ***', '0,27 ***', 'Giảm 78% trầm cảm'],
       ['Tâm sự với bạn (vs không)', '0,37 ***', '0,42 ***', 'Giảm 63% trầm cảm'],
       ['Hỗ trợ gia đình "More" vs hiếm', '0,45 ***', '0,69', 'Giảm 55% trầm cảm'],
       ['Tìm trợ giúp chủ động vs thụ động', '0,48 ***', '0,70', 'Hành vi chủ động'],
       ['Xếp hạng < 60 % (yếu tố nguy cơ)', '1,62 **', '0,98', 'Tăng 62% trầm cảm']],
      widths=[6.0, 3.0, 3.0, 3.5])
H(doc, 'Phản biện', level=2)
P(doc, 'PLOS ONE Q1, Open Access, peer review history công khai. Mẫu LỚN n = 2.716 — đủ power. Đo CẢ 3 trục DASS-21. '
       'PHÁT HIỆN MỚI QUAN TRỌNG NHẤT: kênh giao tiếp gia đình (tâm sự) là yếu tố bảo vệ MẠNH NHẤT — mạnh hơn cả "mức hỗ '
       'trợ" chung. Gợi ý: tập huấn KỸ NĂNG GIAO TIẾP cha–con thay vì chỉ "tăng hỗ trợ" chung chung. Hạn chế: cắt ngang, '
       '1 thành phố Ya An (không đại diện toàn TQ), giới tính p > 0,2 trái với hầu hết NC khác (Wen 2020, Hoa 2024).')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Lặp lại tại VN với mẫu lớn (> 2.000) THPT, đo cả 3 trục DASS-21. Thêm câu hỏi về kênh giao tiếp gia đình. NC dọc '
       '3 năm (như VN21 Trần Thảo Vi). Phát triển module CAN THIỆP KỸ NĂNG GIAO TIẾP cha–con cho đề cương VN.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Mẫu lớn Q1, phát hiện cơ chế bảo vệ then chốt.', bold=True)
save(doc, 'QT47_Dong_PLOS_DASS_2025.docx')

# ============================================================
# BÀI 55 — Chen 2025 COVID Meta 141 NC (abstract-only)
# ============================================================
print('[17/20] Bài 55: Chen 2025 COVID Meta 141 NC')
doc = make_doc()
H(doc, 'Tóm tắt bài QT48 — Meta-analysis 3 cấp về yếu tố nguy cơ và bảo vệ lo âu trẻ em VTN COVID', level=1)
P(doc, '⚠ LƯU Ý: Bản tóm tắt này dựa trên ABSTRACT từ Crossref/Semantic Scholar — chưa có toàn văn (paywall ScienceDirect).',
  bold=True, italic=True)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Protective and risk factors of anxiety in children and adolescents during COVID-19: A systematic review and three '
       'level meta-analysis" của Chen, H., Wang, Q., Zhu, J., Zhu, Y., Yang, F., Hui, J., Tang, X. & Zhang, T. (2025), '
       'đăng trên Journal of Affective Disorders vol. 374, tr. 408–432 (Q1, IF ≈ 6,5) — Elsevier. DOI 10.1016/j.jad.2025.01.029. '
       'Phân tích tổng hợp 141 nghiên cứu, > 1.000.000 trẻ em + VTN trên toàn cầu trong giai đoạn COVID-19.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tổng quan hệ thống + Meta-analysis 3 cấp (three-level meta-analysis). Tổng cộng 141 NC bao phủ HƠN 1 TRIỆU người '
       'tham gia. Phân loại các yếu tố ảnh hưởng theo nhiều domain.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Đã xác định 14 yếu tố BẢO VỆ và 29 yếu tố NGUY CƠ — mất cân đối giữa nguy cơ và bảo vệ. "Chức năng cảm xúc" '
       '(emotional functioning) là yếu tố BẢO VỆ MẠNH. Hỗ trợ gia đình và nguồn lực cộng đồng là các bộ đệm (buffer) '
       'quan trọng. Yếu tố nguy cơ chính: nghiện thiết bị điện tử/internet, mức độ lo âu của người chăm sóc.')
H(doc, 'Phản biện', level=2)
P(doc, 'META 141 NC + 1 TRIỆU trẻ — bằng chứng MẠNH NHẤT về yếu tố nguy cơ/bảo vệ thời COVID. Q1 IF ≈ 6,5. Yếu tố bảo '
       'vệ KÉP: emotional functioning + family support + community resources — phù hợp đề cương VN. "Caregiver anxiety" '
       'là yếu tố nguy cơ → gợi ý can thiệp song song cha mẹ. Internet/device addiction → phù hợp Zheng QT41 (MXH+lo âu) '
       'và Nature SM 2025 (QT27). Mất cân bằng 29 nguy cơ vs 14 bảo vệ → VN cần tăng cường yếu tố bảo vệ. Hạn chế: chưa '
       'có toàn văn → không biết chi tiết effect sizes, OR/RR cụ thể từng yếu tố.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Tải full PDF từ thư viện ĐH để rút trích chi tiết. Lặp lại phân tích cho VN bằng dữ liệu Hoa 2024, Hoàng Trung '
       'Học 2025, Nguyen LX 2023, Hải Phòng 2024, Long An 2025. Can thiệp song song trẻ + cha mẹ.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao (chỉ tính abstract). Cần tải full PDF.', bold=True)
save(doc, 'QT48_Chen_COVID_Meta141_2025_ABSTRACT.docx')

# ============================================================
# BÀI 56 — Zhang 2026 Bayesian MA (abstract-only)
# ============================================================
print('[18/20] Bài 56: Zhang 2026 Bayesian MA')
doc = make_doc()
H(doc, 'Tóm tắt bài QT49 — Bayesian Meta-Analysis CBT trường dài hạn (abstract-only)', level=1)
P(doc, '⚠ LƯU Ý: Bản tóm tắt này dựa trên ABSTRACT từ Crossref — chưa có toàn văn (paywall Wiley).',
  bold=True, italic=True)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Long-Term Effects of School-Based CBT in Low-Risk Children and Adolescents: A Bayesian Meta-Analysis" của Zhang, '
       'X., Liang, Z. & Kang, J. (2026), đăng trên Journal of Clinical Psychology vol. 82, tr. 248–259, online 11/2025, '
       'in print 03/2026 (Q1) — Wiley. DOI 10.1002/jclp.70069. Mẫu: 31 RCT, n = 19.865 trẻ em + VTN nguy cơ thấp — TỔNG '
       'mẫu LỚN NHẤT hiện nay cho can thiệp CBT phổ quát tại trường.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Bayesian Meta-Analysis 31 RCT về CBT phổ quát (universal CBT) tại trường cho trầm cảm và lo âu ở thanh thiếu niên '
       'nguy cơ thấp (low-risk). Theo dõi dài hạn để đánh giá duy trì hiệu quả.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Cải thiện CÓ Ý NGHĨA THỐNG KÊ nhưng KHIÊM TỐN trên triệu chứng trầm cảm; giảm NHỎ trên triệu chứng lo âu. Hiệu '
       'quả DUY TRÌ tới 1 năm sau can thiệp. Kết luận của tác giả: "Chất lượng bằng chứng nền RẤT THẤP khiến phát hiện '
       'này CHƯA ĐỦ vững chắc để hỗ trợ triển khai rộng rãi tại thời điểm này."')
H(doc, 'Phản biện', level=2)
P(doc, 'Bayesian MA 31 RCT, n = 19.865 — tổng mẫu lớn nhất. Q1 Wiley. Phát hiện QUAN TRỌNG: CBT phổ quát có hiệu quả '
       'NHỎ và CHẤT LƯỢNG NỀN THẤP — TRÁI với QT29 BMC NMA 2025 (CBT SUCRA 0,66 hạng 2). Mâu thuẫn này phản ánh khác '
       'biệt giữa CBT PHỔ QUÁT (low-risk, dilution) vs CBT CÓ MỤC TIÊU (high-risk). Phù hợp Brown & Carter 2025 (mindfulness '
       '8.376 HS thất bại). Gợi ý cho VN: nên thử CBT TARGETED (HS có triệu chứng) thay vì universal — giảm dilution.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Tải full PDF để biết SMD cụ thể, heterogeneity I², phân tích phụ. Đề cương VN nên ưu tiên thiết kế TARGETED + '
       'self-referral (theo BESST UK), không universal.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Bayesian MA Q1 — quan trọng cho cảnh báo về universal CBT.', bold=True)
save(doc, 'QT49_Zhang_Bayesian_CBT_2026_ABSTRACT.docx')

# ============================================================
# BÀI 57 — Qiaochu 2025 Mobile CBT SR (abstract-only)
# ============================================================
print('[19/20] Bài 57: Qiaochu 2025 Mobile CBT SR')
doc = make_doc()
H(doc, 'Tóm tắt bài QT50 — Tổng quan hệ thống Mobile CBT cho trẻ em và VTN (abstract-only)', level=1)
P(doc, '⚠ LƯU Ý: Bản tóm tắt này dựa trên ABSTRACT từ Crossref — chưa có toàn văn (paywall Wiley).',
  bold=True, italic=True)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Effectiveness of Mobile-Based Cognitive Behavioural Therapy for Depression and Anxiety in Children and Young People: '
       'A Systematic Review of Randomized Controlled Trials" của Qiaochu, Z. & Wang, Y. (2025), đăng trên Clinical Psychology '
       '& Psychotherapy vol. 32(6), e70173 (Q1) — Wiley. DOI 10.1002/cpp.70173. Tổng quan 9 RCT, N = 2.479 trẻ em + người '
       'trẻ. Đánh giá hiệu quả CBT cung cấp qua MOBILE APP (smartphone).')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tổng quan hệ thống các RCT về Mobile CBT cho trẻ em + người trẻ. 9 RCT đáp ứng tiêu chí được đưa vào tổng hợp. '
       'Phân loại theo outcomes: trầm cảm vs lo âu.')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'PHÁT HIỆN QUAN TRỌNG: 7 trên 8 nghiên cứu (87,5 %) đo TRẦM CẢM cho thấy giảm có ý nghĩa thống kê. NHƯNG chỉ 2 '
       'trên 6 nghiên cứu (33,3 %) đo LO ÂU cho thấy hiệu quả có ý nghĩa. Mobile CBT MẠNH cho trầm cảm, YẾU cho lo âu '
       '(khi không thiết kế đặc thù).')
table(doc, ['Outcome', 'Tỷ lệ NC dương tính', 'Ý nghĩa'],
      [['Trầm cảm', '7/8 (87,5 %)', 'Bằng chứng MẠNH'],
       ['Lo âu', '2/6 (33,3 %)', 'Bằng chứng YẾU']],
      widths=[5.0, 5.0, 5.5])
H(doc, 'Phản biện', level=2)
P(doc, '9 RCT, N = 2.479 — đủ lớn để rút kết luận. Q1 Wiley. Phát hiện cốt lõi: hiệu quả TRẦM CẢM (7/8) nhưng LO ÂU YẾU '
       '(2/6) — Mobile CBT KHÔNG MẠNH cho lo âu khi không thiết kế đặc thù. Mâu thuẫn với Sasaki Japan iCBT (RCT đa trung '
       'tâm dương tính cho SAD) và Walder DMHI MA (g = 0,878 cho SAD-specific). Lý do: các app trong Qiaochu có thể là '
       'cho TRẦM CẢM tổng hợp, không SAD-specific. Hạn chế: chưa có toàn văn → không biết app cụ thể, dropout rates, age '
       'range chi tiết.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Tải full PDF để xem app nào hiệu quả. Cho VN: nếu phát triển app tiếng Việt, nên (a) bắt đầu module trầm cảm '
       '(bằng chứng mạnh), (b) hoặc làm app dành riêng cho LO ÂU XÃ HỘI có hỗ trợ người (theo Walder).')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SR Q1, kết quả khác biệt depression vs anxiety quan trọng.', bold=True)
save(doc, 'QT50_Qiaochu_MobileCBT_2025_ABSTRACT.docx')

# ============================================================
# BÀI 58 — Menon 2025 Scoping LMIC (abstract-only)
# ============================================================
print('[20/20] Bài 58: Menon 2025 Scoping LMIC SEA Pacific')
doc = make_doc()
H(doc, 'Tóm tắt bài QT51 — Scoping Review can thiệp SKTT LMIC Đông Á - Thái Bình Dương', level=1)
P(doc, '⚠ LƯU Ý: Bản tóm tắt này dựa trên ABSTRACT từ Crossref — chưa có toàn văn (paywall SAGE).',
  bold=True, italic=True)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Evaluated Interventions Targeting the Mental Health and Psychosocial Wellbeing of Children and Adolescents: A '
       'Scoping Review Focused on Low- and Middle-Income Countries in East Asia and the Pacific" của Menon, V., Coppard, '
       'M., McEwen, S., Romero, L., Kennedy, E. & Azzopardi, P. (2025), đăng trên Asia Pacific Journal of Public Health '
       'vol. 37(4), tr. 332–346 (Q2) — SAGE. DOI 10.1177/10105395241313154. Scoping Review về can thiệp SKTT + tâm lý xã '
       'hội cho trẻ em và VTN tại các nước thu nhập thấp – trung bình (LMIC) ở Đông Á và Thái Bình Dương — bao gồm Việt Nam.')
H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Scoping review tổng hợp 69 nghiên cứu từ 12 quốc gia LMIC Đông Á + Thái Bình Dương: 32 RCT, 31 nghiên cứu trước–sau, '
       '6 đánh giá sau can thiệp. Phân loại theo loại can thiệp (cá nhân, gia đình, cộng đồng) và đối tượng (phòng ngừa, '
       'điều trị, đáp ứng).')
H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Phát hiện chính: Có sự TẬP TRUNG MẤT CÂN ĐỐI vào "individual capacity" (năng lực cá nhân) — tức can thiệp PHÒNG '
       'NGỪA và "clinical management" (quản lý lâm sàng). KHOẢNG TRỐNG ở: (a) thúc đẩy SK dựa CỘNG ĐỒNG; (b) phòng ngừa '
       'cấp GIA ĐÌNH; (c) dịch vụ ĐÁP ỨNG dài hạn. Phân bố địa lý: phần lớn NC tập trung ở TRUNG QUỐC + Đông Nam Á; các '
       'nước nhỏ hơn và vùng Thái Bình Dương có ĐẠI DIỆN TỐI THIỂU.')
H(doc, 'Phản biện', level=2)
P(doc, '69 NC / 12 nước LMIC Đông Á + TBD — bao gồm Việt Nam. Phản ánh KHU VỰC trực tiếp cho đề tài. KHOẢNG TRỐNG xác '
       'định rõ: thiếu can thiệp CỘNG ĐỒNG + GIA ĐÌNH + DỊCH VỤ DÀI HẠN. Phù hợp với phát hiện Dong QT47 (kênh giao tiếp '
       'gia đình OR = 0,22) và UNICEF VN22 + đề cương VN. Cảnh báo: TQ + ĐNA chiếm phần lớn → có thể thiếu chú trọng VN '
       'cụ thể. Hạn chế: chưa có toàn văn → chưa biết VN có bao nhiêu NC, loại nào.')
H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Tải full PDF để xem chi tiết VN trong scoping. Đề cương VN giai đoạn 2 nên ƯU TIÊN can thiệp CỘNG ĐỒNG + GIA ĐÌNH '
       'thay vì chỉ trường — sẽ lấp khoảng trống. Module GIAO TIẾP GIA ĐÌNH (theo Dong QT47) là khả thi.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Scoping review khu vực ĐÁ-TBD LMIC — trực tiếp cho VN.', bold=True)
save(doc, 'QT51_Menon_LMIC_SEA_Pacific_2025_ABSTRACT.docx')

print('\nDONE 20/20. All summaries saved to Tom-tat-tung-bai/')
