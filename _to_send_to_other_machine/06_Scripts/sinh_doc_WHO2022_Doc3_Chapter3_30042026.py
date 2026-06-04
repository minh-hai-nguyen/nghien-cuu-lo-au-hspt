# -*- coding: utf-8 -*-
"""WHO 2022 World Mental Health Report — Doc 3 — Chương 3 chi tiết.
Dịch FULL 1-1 PDF idx 57-90 (trang 35-68 sách) — World mental health today.
Workflow chuẩn: tiếng Việt thuần, page markers cam, từ viết tắt lần đầu đỏ đậm,
bảng Word tái tạo, hình từ PDF, phản biện cuối + áp dụng VN.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\WHO_2022_Doc3_Chapter3_dich_phan_bien_30042026.docx'
FIG_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\02_Papers-goc\GBD_WHO\figures'

RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)
ORANGE = RGBColor(0xE0, 0x6C, 0x00); RED_BOLD = RGBColor(0xCC, 0, 0)
DRED = RGBColor(0x9B, 0x1B, 0x1B)
BOX_BG = 'FFF2CC'  # light yellow for Box callouts

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(12)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)

def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
    d.add_paragraph()

def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def page_marker(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'— — — {text} — — —')
    r.bold=True; r.italic=True; r.font.size=Pt(11)
    r.font.color.rgb=ORANGE; r.font.name='Times New Roman'
def acro(viet, eng, abbrev):
    """First-time abbreviation: dark red bold + inline gloss"""
    p=d.add_paragraph()
    r=p.add_run(f'{viet} ({eng}'); r.bold=True; r.font.color.rgb=DRED; r.font.size=Pt(12); r.font.name='Times New Roman'
    if abbrev:
        r2=p.add_run(f' — {abbrev}'); r2.bold=True; r2.font.color.rgb=DRED; r2.font.size=Pt(12); r2.font.name='Times New Roman'
    rclose=p.add_run(')'); rclose.bold=True; rclose.font.color.rgb=DRED; rclose.font.size=Pt(12); rclose.font.name='Times New Roman'
def crit_para(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run(text); r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
def insert_image(fname, caption, width_cm=15):
    path = os.path.join(FIG_DIR, fname)
    if not os.path.exists(path):
        nr(f'[Hình không tải được: {fname}]', italic=True, color=GRAY, size=11); return
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.add_picture(path, width=Cm(width_cm))
    cap = d.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rcap = cap.add_run(caption)
    rcap.italic = True; rcap.bold = True; rcap.font.size = Pt(10)
    rcap.font.color.rgb = ORANGE; rcap.font.name = 'Times New Roman'

def box_open(box_id, box_title, box_type='INSIGHT'):
    """Open a Box (BOX 3.x) — bordered/shaded paragraph block"""
    p = d.add_paragraph()
    r = p.add_run(f'[{box_type}] {box_id} — {box_title}')
    r.bold = True; r.font.size = Pt(12); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
    # Add yellow background to indicate box content
def box_para(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(text); r.font.size = Pt(12); r.font.name = 'Times New Roman'
def box_close():
    nr('— — — Hết Box — — —', italic=True, color=GRAY, size=10)

def narrative_open(name, country, story_title):
    p = d.add_paragraph()
    r = p.add_run(f'[CÂU CHUYỆN] "{story_title}" — {name} ({country})')
    r.bold = True; r.italic = True; r.font.size = Pt(12); r.font.color.rgb = GREEN
    r.font.name = 'Times New Roman'
def narrative_para(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.7)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    r.font.color.rgb = GRAY

# ========== TRANG BÌA ==========
title("BẢN DỊCH + PHẢN BIỆN CHI TIẾT", 14)
title("BÁO CÁO SỨC KHOẺ TÂM THẦN THẾ GIỚI 2022", 14)
title("CHƯƠNG 3 — SỨC KHOẺ TÂM THẦN THẾ GIỚI HÔM NAY", 16)
nr("")
subtitle("Chapter 3 — World mental health today")
subtitle("Dịch tễ — Hệ quả kinh tế — Khoảng trống chính — Nhu cầu chăm sóc")
subtitle("(Epidemiology — Economic costs — Key gaps — Demand for care)")
nr("")
title("Doc 3 trong bộ 9 doc", 12)
subtitle("Doc 0 (Index + Foreword + Exec Summary đã có) | Doc 3 (đang đọc) | Doc 7, 2, 6, 5, 4, 1, 8 (sẽ làm sau)")
nr("")
nr("Tổ chức Y tế Thế giới (World Health Organization — WHO)", bold=True)
nr("Geneva, 17/06/2022 — 296 trang — Open Access (CC BY-NC-SA 3.0 IGO)", size=11)
nr("ISBN 978-92-4-004933-8 (electronic) / 978-92-4-004934-5 (print)", size=11, color=GRAY)
nr("")
nr("PDF gốc: 02_Papers-goc/GBD_WHO/WHO_2022_World_Mental_Health_Report.pdf (3,6 MB / 296 trang)",
   italic=True, size=10, color=GRAY)
nr("Phạm vi Doc 3: PDF idx 57-90 = trang 35-68 sách = Chương 3 đầy đủ (~33 trang)",
   italic=True, size=10, color=GRAY)
nr("Trợ lý nghiên cứu — 30/04/2026 — tiếng Việt thuần — chú thích Anh trong ngoặc",
   italic=True, size=10, color=GRAY)

# ========== THÔNG TIN THƯ MỤC ==========
H1("THÔNG TIN THƯ MỤC CHƯƠNG 3")
tbl(['Mục', 'Nội dung'], [
    ['Tên chương (tiếng Anh)', 'Chapter 3 — World mental health today'],
    ['Tên chương (tiếng Việt)', 'Chương 3 — Sức khoẻ tâm thần thế giới hôm nay'],
    ['Báo cáo cha', 'WHO 2022 World Mental Health Report — Transforming mental health for all'],
    ['Tổ chức', 'World Health Organization (WHO)'],
    ['Năm xuất bản', '2022 (17/06/2022)'],
    ['Trang sách', '35-68 (PDF idx 57-90)'],
    ['Số trang chương', '34 trang'],
    ['Số mục lớn', '4 mục: 3.1 Dịch tễ, 3.2 Hệ quả kinh tế, 3.3 Khoảng trống, 3.4 Rào cản'],
    ['Box info', '3 BOX: Box 3.1 Data assessing, Box 3.2 COVID-19, Box 3.3 Severity vertical equity'],
    ['Hình minh hoạ', '13 FIG (3.1 đến 3.13) — đã trích PNG đầy đủ'],
    ['Câu chuyện cá nhân', '4 NARRATIVE: Marie (Cameroon), Eleni (Ethiopia), Mrs BN (anonymous), Steven (UK), Odireleng (Botswana)'],
    ['Số tham khảo trích', '~60 tham khảo (số 96 đến 158+)'],
    ['DOI / Trích dẫn', 'World mental health report. Geneva: WHO; 2022. Chapter 3, pp 35-68.'],
    ['URL', 'https://www.who.int/publications/i/item/9789240049338'],
    ['Trạng thái mở', 'Open Access dưới giấy phép CC BY-NC-SA 3.0 IGO'],
])

# ========== BẢNG TỪ VIẾT TẮT ==========
H1("BẢNG TỪ VIẾT TẮT DÙNG TRONG CHƯƠNG 3")
nr("Lần đầu mỗi từ viết tắt xuất hiện trong văn bản dưới đây sẽ được tô MÀU ĐỎ ĐẬM + chú thích inline.")
tbl(['Tiếng Việt', 'Tiếng Anh đầy đủ', 'Viết tắt'], [
    ['Tổ chức Y tế Thế giới', 'World Health Organization', 'WHO'],
    ['Nước thu nhập thấp + trung bình', 'Low- and Middle-Income Countries', 'LMICs'],
    ['Năm sống điều chỉnh theo khuyết tật', 'Disability-Adjusted Life-Year', 'DALY'],
    ['Năm sống với khuyết tật', 'Years Lived with Disability', 'YLDs'],
    ['Năm sống bị mất do tử vong sớm', 'Years of Life Lost', 'YLLs'],
    ['Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019', 'Global Burden of Diseases, Injuries and Risk Factors Study 2019', 'GBD 2019'],
    ['Ước tính Sức khoẻ Toàn cầu của WHO', 'WHO Global Health Estimates', 'GHE'],
    ['Viện Đo lường + Đánh giá Y tế', 'Institute of Health Metrics and Evaluation', 'IHME'],
    ['Phân loại bệnh quốc tế phiên bản 11', 'International Classification of Diseases, 11th Revision', 'ICD-11'],
    ['Sổ tay chẩn đoán + thống kê rối loạn tâm thần phiên bản 5', 'Diagnostic and Statistical Manual of Mental Disorders, 5th edition', 'DSM-5'],
    ['Bệnh không lây nhiễm', 'Noncommunicable Disease', 'NCD'],
    ['Rối loạn tăng động giảm chú ý', 'Attention-Deficit Hyperactivity Disorder', 'ADHD'],
    ['Rối loạn phổ tự kỷ', 'Autism Spectrum Disorder', 'ASD'],
    ['Rối loạn căng thẳng sau sang chấn', 'Post-Traumatic Stress Disorder', 'PTSD'],
    ['Rối loạn trầm cảm chính', 'Major Depressive Disorder', 'MDD'],
    ['Khuyết tật trí tuệ phát triển không rõ nguyên nhân', 'Idiopathic Developmental Intellectual Disability', 'IID'],
    ['Mục tiêu Phát triển Bền vững', 'Sustainable Development Goal', 'SDG'],
    ['Bao phủ Y tế Toàn dân', 'Universal Health Coverage', 'UHC'],
    ['Sức khoẻ ban đầu', 'Primary Health Care', 'PHC'],
    ['Tổng sản phẩm quốc nội', 'Gross Domestic Product', 'GDP'],
    ['Bệnh lây qua đường tình dục', 'Sexually Transmitted Infections', 'STIs'],
    ['Vi-rút gây suy giảm miễn dịch ở người', 'Human Immunodeficiency Virus / Acquired Immunodeficiency Syndrome', 'HIV/AIDS'],
])

# ========== PHẦN A — TRANG MỞ CHƯƠNG (PDF idx 57) ==========
H1("PHẦN A — TRANG MỞ CHƯƠNG 3")
page_marker('Trang 35 sách (PDF idx 57), WHO 2022 — Trang mở chương')
nr('CHƯƠNG 3', bold=True)
nr('CHỦ ĐỀ CHÍNH: DỊCH TỄ HỌC — HỆ QUẢ KINH TẾ — KHOẢNG TRỐNG CHÍNH — NHU CẦU CHĂM SÓC',
   bold=True, color=BLUE)
nr('SỨC KHOẺ TÂM THẦN THẾ GIỚI HÔM NAY', bold=True, size=14, color=BLUE)
subtitle('(EPIDEMIOLOGY — ECONOMIC COSTS — KEY GAPS — DEMAND FOR CARE — World mental health today)')

# ========== PHẦN B — TÓM TẮT CHƯƠNG (PDF idx 58) ==========
H1("PHẦN B — TÓM TẮT CHƯƠNG (Chapter summary)")
page_marker('Trang 36 sách (PDF idx 58), WHO 2022 — Chapter summary')

H2("Tóm tắt chương")
nr("Trong chương này, chúng tôi phác thảo trạng thái sức khoẻ tâm thần (SKTT) và các hệ thống "
   "SKTT trên thế giới, và cho thấy nhu cầu SKTT là CAO trong khi các đáp ứng KHÔNG ĐỦ và KHÔNG "
   "PHÙ HỢP. Chúng tôi trình bày dữ liệu mới nhất hiện có về tỉ lệ phổ biến và chi phí toàn cầu "
   "của các rối loạn tâm thần — nhìn xa hơn tác động của tử vong và khuyết tật để xét đến cả "
   "các chi phí kinh tế và xã hội KHỔNG LỒ liên quan. Chương này cũng nhấn mạnh kết quả của "
   "Bản đồ Sức khoẻ Tâm thần (Mental Health Atlas) gần nhất của ")
acro('Tổ chức Y tế Thế giới', 'World Health Organization', 'WHO')
nr("để hé lộ một số khoảng trống và rào cản nghiêm trọng + dai dẳng trong chăm sóc SKTT trên thế giới.")

H2("Thông điệp chính của chương")
nr("• Ở MỌI quốc gia, các rối loạn tâm thần có tỉ lệ phổ biến CAO và phần lớn KHÔNG ĐƯỢC ĐIỀU TRỊ.")
nr("• Các rối loạn tâm thần là nguyên nhân HÀNG ĐẦU của số năm sống với khuyết tật — và tự tử "
   "vẫn là nguyên nhân tử vong CHÍNH trên toàn cầu.")
nr("• Hệ quả kinh tế của các tình trạng SKTT là KHỔNG LỒ — với tổn thất năng suất CAO HƠN ĐÁNG "
   "KỂ so với chi phí chăm sóc trực tiếp.")
nr("• Các hệ thống SKTT trên toàn thế giới đều có những khoảng trống LỚN về quản trị, nguồn lực, "
   "dịch vụ, thông tin và công nghệ cho SKTT.")
nr("• Một số yếu tố ngăn cản người dân tìm kiếm sự giúp đỡ cho các tình trạng SKTT — bao gồm "
   "tiếp cận DỊCH VỤ chất lượng còn HẠN CHẾ, mức độ HIỂU BIẾT về SKTT THẤP, và KỲ THỊ phổ biến.")

# ========== PHẦN C — MỞ ĐẦU CHƯƠNG (PDF idx 59) ==========
H1("PHẦN C — MỞ ĐẦU CHƯƠNG (3.0)")
page_marker('Trang 37 sách (PDF idx 59), WHO 2022 — Section opener + FIG 3.1')

nr("Mặc dù SKTT có tầm quan trọng SỐNG CÒN đối với sức khoẻ và hạnh phúc của chúng ta, quá NHIỀU "
   "người trong số chúng ta KHÔNG nhận được sự hỗ trợ mình cần. Năm 2019, ước tính có MỘT TRONG "
   "TÁM người trên thế giới đang sống chung với một rối loạn tâm thần (96). Cùng lúc đó, các "
   "dịch vụ, kỹ năng và kinh phí dành cho SKTT vẫn THIẾU HỤT, và CÒN XA mức cần — đặc biệt ở các ")
acro('Nước thu nhập thấp + trung bình', 'Low- and Middle-Income Countries', 'LMICs')
nr(".")

nr("Ở MỌI quốc gia, các tình trạng SKTT phổ biến RỘNG (nhưng bị HIỂU SAI) và KHÔNG ĐƯỢC ĐIỀU TRỊ "
   "đủ — và các dịch vụ giải quyết chúng KHÔNG ĐƯỢC cấp đủ nguồn lực (xem Hình 3.1). Và như đã "
   "thảo luận ở Chương 2 \"Nguyên tắc và động lực trong SKTT công cộng\", các yếu tố sinh học - "
   "tâm lý - xã hội tương tác làm suy yếu SKTT — từ căng thẳng cấp cộng đồng như nghèo đói, xung "
   "đột và bất bình đẳng xã hội đến các yếu tố cá nhân như giá trị bản thân thấp — sẽ TIẾP TỤC "
   "tạo ra mối đe doạ với SKTT trong tương lai gần.")

nr("Chương này trình bày dữ liệu mới nhất có sẵn tại thời điểm viết (xem Box 3.1 \"Dữ liệu để "
   "đánh giá SKTT thế giới\"). Trong hầu hết các trường hợp, dữ liệu CÓ TRƯỚC đại dịch COVID-19 "
   "— vốn đã làm TRẦM TRỌNG đáng kể các yếu tố nguy cơ với các tình trạng SKTT cho nhiều người. "
   "Đại dịch CHẮC CHẮN sẽ tác động đến tỉ lệ phổ biến và gánh nặng của các rối loạn tâm thần — "
   "đồng thời tiếp cận với dịch vụ SKTT cũng đã bị ảnh hưởng (xem Chương 2, In focus: COVID-19 "
   "và SKTT). Một sự gia tăng dài hạn về số lượng và mức độ nặng của các tình trạng SKTT trên "
   "toàn thế giới đã được DỰ ĐOÁN — và như được trình bày dưới đây, các ước tính toàn cầu mới "
   "nhất đã XÁC NHẬN điều này (97, 69).")

insert_image('WHO2022_Fig3_1_service_gap.png',
             'Hình 3.1 (trang 37 sách) — "Các tình trạng SKTT phổ biến rộng, không được điều trị '
             'và không được cấp nguồn lực." Nguồn: IHME 2019 (98); WHO 2021 (5). '
             'PHỔ BIẾN: 1/8 người sống với một tình trạng SKTT. KHÔNG ĐIỀU TRỊ: 71% người có '
             'loạn thần KHÔNG nhận được dịch vụ SKTT. KHÔNG NGUỒN LỰC: chỉ 2% ngân sách y tế '
             'trung bình dành cho SKTT.', width_cm=14)

# ========== PHẦN D — BOX 3.1 (PDF idx 60) ==========
H1("PHẦN D — BOX 3.1: DỮ LIỆU ĐỂ ĐÁNH GIÁ SKTT THẾ GIỚI")
page_marker('Trang 38 sách (PDF idx 60), WHO 2022 — BOX 3.1 INSIGHT')

box_open('BOX 3.1', 'Dữ liệu để đánh giá SKTT thế giới (Data for assessing world mental health)', 'INSIGHT')
box_para("Để giao tiếp với phổ rộng nhất các bên liên quan, báo cáo này thường dùng thuật ngữ ô "
         "che \"các tình trạng SKTT\" (mental health conditions) — bao gồm: rối loạn tâm thần "
         "(mental disorders), khuyết tật tâm lý-xã hội (psychosocial disabilities) và các trạng "
         "thái tâm thần khác liên quan đến đau khổ đáng kể, suy giảm chức năng hoặc nguy cơ tự "
         "hại.")
box_para("Tuy nhiên, khi mô tả TỈ LỆ PHỔ BIẾN và các ước tính toàn cầu trong chương này, chúng "
         "tôi dùng \"rối loạn tâm thần\" — vì thuật ngữ này phản ánh chính xác hơn dữ liệu đang "
         "được thu thập + báo cáo, và phạm vi được xác định rõ trong ICD-11 của WHO. Tương tự, "
         "trong chương này chúng tôi dùng các phân loại chẩn đoán như \"rối loạn trầm cảm\" hay "
         "\"rối loạn lo âu\" — thay vì các thuật ngữ rộng hơn \"trầm cảm\" và \"lo âu\" như ở "
         "các chương khác.")
box_para("Các rối loạn tâm thần KHÁC BIỆT với rối loạn thần kinh và rối loạn sử dụng chất. Hai "
         "nhóm sau — dù không phải trọng tâm của báo cáo — vẫn được nhắc trong chương này để mô "
         "tả phạm vi nhu cầu mà người ra quyết định SKTT thường phải chịu trách nhiệm ở các LMICs.")
box_para("Đo lường + giám sát tỉ lệ mới mắc, hiện mắc và tử vong cũng như phân bố bệnh tật và các "
         "yếu tố quyết định trong và giữa các quần thể — đặc trưng định nghĩa của dịch tễ học — "
         "cung cấp THÔNG TIN SỐNG CÒN cho lập kế hoạch + cung cấp + đánh giá dịch vụ y tế. Nguồn "
         "dữ liệu dịch tễ học quốc tế chính được dùng trong chương này là ")
box_para("(a) Ước tính Sức khoẻ Toàn cầu của WHO (WHO Global Health Estimates — GHE); và (b) "
         "Nghiên cứu Gánh nặng Bệnh tật, Chấn thương và Yếu tố Nguy cơ Toàn cầu 2019 (GBD 2019) "
         "của Viện Đo lường + Đánh giá Y tế (IHME). Hai nguồn này LIÊN KẾT CHẶT về các ước tính "
         "SKTT. Cùng nhau, chúng cung cấp tỉ lệ phổ biến điểm + ước tính gánh nặng bệnh tật cho "
         "tất cả các loại bệnh truyền nhiễm + không truyền nhiễm + chấn thương.")
box_para("Thuật ngữ \"gánh nặng bệnh tật\" CHỈ được dùng liên quan đến các đánh giá dịch tễ đã "
         "công bố. Đây là thuật ngữ chuẩn dùng trong y tế công cộng cho các ước tính tác động "
         "cấp dân số (vd: DALY, YLL, YLD).")
box_para("Các ước tính cần được giải thích THẬN TRỌNG. Các nghiên cứu toàn cầu lâu năm này cung "
         "cấp BẰNG CHỨNG TỐT NHẤT có sẵn về phạm vi, phân bố và tác động sức khoẻ công cộng của "
         "các rối loạn tâm thần — phân theo nhóm tuổi, giới tính và quốc gia (gộp theo thu nhập "
         "hoặc địa lý). Tuy nhiên dù các ước tính hiện tại đã tích hợp dữ liệu mới nhất + tiến "
         "bộ phương pháp mô hình hoá, chúng vẫn KHÔNG CHẮC CHẮN do thiếu hụt dữ liệu dịch tễ "
         "học cho rối loạn tâm thần ở nhiều quốc gia. Đặc biệt, các ước tính thường dựa trên dữ "
         "liệu đầu vào KHÔNG ĐẦY ĐỦ — không bao quát mọi tham số hoặc mọi quốc gia — và dựa trên "
         "thông tin LẠC HẬU hoặc CHẤT LƯỢNG KÉM. Hơn nữa, cần nhận ra rằng các rối loạn tâm thần "
         "có thể được khái niệm hoá theo nhiều cách KHÁC NHAU giữa các nền văn hoá — gây thách "
         "thức cho việc đo lường chúng từ một điểm tham chiếu cụ thể như trong các nghiên cứu "
         "gánh nặng bệnh tật toàn cầu.")
box_close()

OUT_DIR = os.path.dirname(OUT)
os.makedirs(OUT_DIR, exist_ok=True)
d.save(OUT)
print(f"Wrote (Part 1/3): {OUT}")
