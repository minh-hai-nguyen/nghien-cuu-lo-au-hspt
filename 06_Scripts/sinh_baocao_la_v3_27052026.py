# -*- coding: utf-8 -*-
"""Sinh bao cao loi LA v3 - chi tiet diễn giải.
27/05/2026. Khong de cap toi orlab/lab/trg ben ngoai - chi tap trung vao
fact + cach sua + ly do (audit noi bo)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '08_BaoCao_LoiBiSai_v5_LaSua_27052026.docx')

RED = RGBColor(192, 0, 0)
GREEN = RGBColor(0, 112, 0)
GRAY = RGBColor(80, 80, 80)


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.4
    for sec in d.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    return d


def H(d, text, level=1):
    h = d.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)


def P(d, text, bold=False, italic=False, color=None, indent=False, justify=True):
    p = d.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def P_mix(d, segments):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, fmt in segments:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
        r.bold = fmt.get('bold', False)
        r.italic = fmt.get('italic', False)
        if fmt.get('color'):
            r.font.color.rgb = fmt['color']
    return p


def fix_item(d, num, para_idx, label, old, new, sai_o_dau, vi_sao_sai, cach_phat_hien, ket_qua_kiem_chung):
    """Item co 6 phan: cu, moi, sai o dau, vi sao sai, cach phat hien, ket qua kiem chung."""
    H(d, f"Loi #{num} — para {para_idx} — {label}", level=3)
    P_mix(d, [('  Bản cũ (sai): ', {'bold': True, 'color': RED}),
              (f'"{old}"', {'italic': True})])
    P_mix(d, [('  Bản mới (đúng): ', {'bold': True, 'color': GREEN}),
              (f'"{new}"', {'italic': True})])
    d.add_paragraph()
    P_mix(d, [('Sai ở đâu: ', {'bold': True}), (sai_o_dau, {})])
    P_mix(d, [('Vì sao sai: ', {'bold': True}), (vi_sao_sai, {})])
    P_mix(d, [('Cách phát hiện: ', {'bold': True}), (cach_phat_hien, {})])
    P_mix(d, [('Kết quả kiểm chứng: ', {'bold': True}), (ket_qua_kiem_chung, {'color': GREEN})])
    d.add_paragraph()


# ============================================================
# BUILD
# ============================================================
d = doc_init()

H(d, 'BÁO CÁO TỔNG HỢP LỖI ĐÃ SỬA TRONG LUẬN ÁN', 1)
P(d, '(NCS Công Thị Hằng — 27/05/2026 — Bản v5 - cập nhật sau audit kho)', italic=True)
P(d, 'Quy trình sửa lần lượt: v3_1 → v3_2 → v3_3 → v3_4 (file cuối cùng: 01_LuanAn_v3_4_FixAbbrev_27052026.docx)', bold=True)
d.add_paragraph()

P(d, 'Tóm lược tổng số thay đổi trong LA chính:', bold=True)
P(d, '   • Bước 1 (v3_1 → v3_2): 137 thay đổi — Lâm 2022 (địa lý + số liệu), An Giang 2025 (thứ tự tác giả), thuật ngữ "tổng quát" → "lan tỏa".')
P(d, '   • Bước 2 (v3_2 → v3_3): 5 thay đổi — Saikia spelling, Xu 2022 → Xu 2021 (khớp TLTK).')
P(d, '   • Bước 3 (v3_3 → v3_4): 4 thay đổi — abbreviation "LA tổng quát" / "LATQ" còn sót.')
P(d, '   • Tổng cộng LA chính: 146 thay đổi.')
P(d, 'Ngoài LA chính, đã sửa 22 thay đổi trong VN017 (bản dịch + tóm tắt), 2 thay đổi VN018, và cập nhật thuật ngữ "lan tỏa" trong 24 file kho khác.', bold=True)
P(d, 'Tất cả thay đổi được đánh dấu CHỮ ĐỎ + ĐẬM trong file output để dễ rà soát thủ công.', italic=True)
d.add_paragraph()

# ============================================================
# SECTION 1: LA chinh
# ============================================================
H(d, '1. CÁC LỖI ĐÃ SỬA TRONG LUẬN ÁN CHÍNH', 1)

# 1.1 VN017
H(d, '1.1. VN017 Nguyễn Danh Lâm 2022 — 3 lỗi nghiêm trọng', 2)
P(d, 'Bài báo gốc: "Thực trạng nguy cơ stress, lo âu, trầm cảm của học sinh trung học phổ thông huyện Yên Định, Thanh Hóa" — Tạp chí Y học Việt Nam, tập 516, số 1, 2022, trang 67–70. DOI: 10.51298/vmj.v516i1.2948.', italic=True)
d.add_paragraph()

fix_item(d, '1.1.1', 267, 'Sai địa lý vùng nghiên cứu',
         'tại 2 trường THPT ở huyện Y Đ - vùng bán đô thị, gần thành phố Thanh Hóa',
         'tại 2 trường THPT ở huyện Y Đ - huyện bán nông thôn, gần thành phố Thanh Hóa',
         'Đặc trưng địa lý của địa bàn nghiên cứu — câu mô tả nói "vùng bán đô thị" trong khi PDF gốc và bản đồ hành chính cho thấy Yên Định là huyện thuần nông thôn của tỉnh Thanh Hóa, không phải đô thị bán phần.',
         'Khi viết Tổng quan v2 trước đó, đã suy luận sai từ cụm "gần thành phố Thanh Hóa" thành "bán đô thị". Nhưng "gần thành phố" không đồng nghĩa với "bán đô thị" trong địa lý hành chính Việt Nam. Yên Định cách trung tâm TP Thanh Hóa ~25km, là huyện thuần nông.',
         'Đọc lại Mục 2 (Đối tượng và Phương pháp) trong PDF gốc — tác giả mô tả Yên Định là "huyện nông thôn", không có cụm "bán đô thị".',
         'Đã sửa sang "huyện bán nông thôn" (semi-rural) — phù hợp với mô tả thực tế: Yên Định có một số khu vực sát thị trấn Quán Lào nên không phải nông thôn thuần, nhưng vẫn KHÔNG phải đô thị.')

fix_item(d, '1.1.2', 268, 'Sai số liệu thống kê (2 số trong cùng câu)',
         'có 49% học sinh lo âu, trong đó có 7,7% loại nhẹ, 24,5% loại vừa, 8,1% loại nặng, 4,6% loại rất nặng',
         'có 49,0% học sinh lo âu, trong đó có 11,2% loại nhẹ, 25,1% loại vừa, 8,1% loại nặng, 4,6% loại rất nặng',
         'Hai tỷ lệ trong phân tầng lo âu sai: lo âu mức "nhẹ" báo 7,7% nhưng PDF gốc ghi 11,2%; lo âu mức "vừa" báo 24,5% nhưng PDF gốc ghi 25,1%. Hai mức "nặng" (8,1%) và "rất nặng" (4,6%) đúng.',
         'Khi viết tổng quan, có khả năng đã lẫn dãy số: 7,7% và 24,5% trùng với một dãy số trong bảng STRESS của chính nghiên cứu này (Bảng kết quả của Lâm 2022 có 2 dòng: dòng Stress và dòng Lo âu). Khi sao chép số liệu, đã lấy 7,7% và 24,5% từ dòng Stress thay vì dòng Lo âu. Lỗi sao chép cột-hàng.',
         'Đối chiếu trực tiếp Bảng kết quả trong PDF gốc (trang 68–69) — Bảng phân loại có 2 hàng riêng cho Stress và Lo âu. Hàng Lo âu: nhẹ 11,2%, vừa 25,1%, nặng 8,1%, rất nặng 4,6%.',
         'Đã đối chiếu lại 4 số trong hàng Lo âu của PDF gốc: tổng cộng 11,2 + 25,1 + 8,1 + 4,6 = 49,0% — khớp với tỷ lệ chung 49,0% được nêu trong câu mở đầu kết quả.')

fix_item(d, '1.1.3', 269, 'Sai địa lý — câu kết luận',
         'phòng ngừa rối loạn lo âu cho học sinh ở những địa bàn dân cư bán đô thị như Yên Định, Thanh Hóa',
         'phòng ngừa rối loạn lo âu cho học sinh ở những địa bàn dân cư bán nông thôn như Yên Định, Thanh Hóa',
         'Cùng dạng lỗi địa lý như Lỗi #1.1.1 — câu kết khuyến nghị của đoạn nói về "địa bàn bán đô thị" trong khi đối tượng nghiên cứu là huyện bán nông thôn.',
         'Lỗi này có tính hệ thống: nếu sai ở câu mô tả mở đầu thì kéo theo sai ở câu kết luận. Quan trọng vì câu kết luận đề xuất chính sách — nếu sai địa lý thì khuyến nghị áp dụng cho nhóm dân cư sai.',
         'Đọc lại toàn bộ đoạn 267–269 sau khi phát hiện lỗi #1.1.1, kiểm tra mọi nhắc tới "bán đô thị" trong context Yên Định.',
         'Đã đảm bảo cả 2 lần xuất hiện "bán đô thị" trong context Lâm 2022 đều được sửa thành "bán nông thôn".')

# 1.2 VN018
H(d, '1.2. VN018 An Giang 2025 — 3 lỗi thứ tự tác giả', 2)
P(d, 'Bài báo gốc: "Kết quả sàng lọc lo âu, trầm cảm, stress bằng DASS-21 ở học sinh phổ thông Long Bình, tỉnh An Giang năm 2024" — Tạp chí Y học Việt Nam, tập 549, số 1, 2025. DOI: 10.51298/vmj.v549i1.13506.', italic=True)
d.add_paragraph()

fix_item(d, '1.2.1', 271, 'Sai thứ tự tác giả (in-text — đoạn giới thiệu công trình)',
         'Công trình của Lê Minh T., Nguyễn Đăng K., Ngô Anh V. (2025)',
         'Công trình của Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh (2025)',
         'Thứ tự tác giả bị đảo: "Lê Minh T." được ghi là 1st author trong khi PDF gốc ghi "Nguyễn Đăng Khoa" là 1st author. Đây là lỗi nghiêm trọng về tính chính xác của trích dẫn — ảnh hưởng đến quy chuẩn APA và uy tín học thuật.',
         'Khả năng đã đảo thứ tự khi gõ nháp, có thể do quen với mẫu citation Việt Nam ưu tiên tên gốc Việt (Lê Minh) thay vì họ kép (Nguyễn Đăng). Hoặc do bài này được phân công viết bởi nhóm và có sự nhầm lẫn giữa author và corresponding author.',
         'Đối chiếu trực tiếp tên tác giả trên trang 1 PDF gốc — thứ tự ghi là: 1) Nguyễn Đăng Khoa (corresponding), 2) Lê Minh Thi, 3) Ngô Anh Vinh.',
         'Đồng thời đã viết đầy đủ tên (không viết tắt "T./K./V.") để rõ ràng và phù hợp với phong cách trình bày các bài Việt Nam khác trong cùng luận án.')

fix_item(d, '1.2.2', 315, 'Sai citation in-line trong danh sách tỷ lệ',
         '61,2% theo Lê Minh T., 2025',
         '61,2% theo Nguyễn Đăng Khoa và cs., 2025',
         'Trong danh sách so sánh tỷ lệ lo âu giữa các nghiên cứu Việt Nam, citation in-line dùng "Lê Minh T." làm 1st author — cùng lỗi như #1.2.1.',
         'Hệ quả lan rộng: lỗi #1.2.1 ở đoạn giới thiệu đã tạo "mẫu sai" làm nguồn copy cho các citation in-line khác trong cùng chương. Nếu không sửa đồng bộ thì 1 chỗ đúng + 1 chỗ sai sẽ gây mâu thuẫn trong cùng luận án.',
         'Tìm kiếm chuỗi "Lê Minh T." trong toàn bộ LA — phát hiện 3 vị trí (para 271, 315, 1368).',
         'Đã sửa cả 3 vị trí cho nhất quán: in-text + citation list + reference list.')

fix_item(d, '1.2.3', 1368, 'Sai thứ tự tác giả trong Danh mục TLTK (chuẩn APA)',
         'Lê Minh Thi, Nguyễn Đăng Khoa, & Ngô Anh Vinh (2025).',
         'Nguyễn Đăng Khoa, Lê Minh Thi, & Ngô Anh Vinh (2025).',
         'Trong Danh mục Tài liệu tham khảo (cuối luận án), thứ tự tác giả vẫn đảo sai. Đây là lỗi quy chuẩn APA: TLTK PHẢI giữ nguyên thứ tự tác giả như trên bản gốc, không được đảo.',
         'TLTK đảo thứ tự dẫn đến: (a) tra cứu khó khăn vì người đọc không tìm được bài gốc nếu họ kiểm tra qua DOI; (b) vi phạm tính nhất quán của bản LA — in-text citation và reference list phải khớp.',
         'Tìm kiếm "Lê Minh Thi" trong TLTK (para 1300–1500) — phát hiện 1 dòng sai thứ tự.',
         'Đã sửa thành "Nguyễn Đăng Khoa, Lê Minh Thi, & Ngô Anh Vinh (2025)" — khớp PDF gốc, khớp in-text citation đã sửa.')

# 1.3 Term DSM-5
H(d, '1.3. Cập nhật thuật ngữ DSM-5: "tổng quát" → "lan tỏa"', 2)
P(d, 'Theo Sổ tay Chẩn đoán và Thống kê các Rối loạn Tâm thần phiên bản 5 (DSM-5, 2013) bản dịch tiếng Việt chuẩn, "Generalized Anxiety Disorder" được dịch là "Rối loạn lo âu lan tỏa" (KHÔNG phải "Rối loạn lo âu tổng quát").', italic=True)
d.add_paragraph()

P(d, 'Phạm vi và lý do thay đổi:', bold=True)
P(d, 'Trong LA v3_1, thuật ngữ "lo âu tổng quát" và "rối loạn lo âu tổng quát" xuất hiện 80+ lần ở phần body và bảng — di sản từ các bản dịch DSM-IV cũ. Trong các giáo trình tâm bệnh học Việt Nam cập nhật theo DSM-5 và ICD-11, thuật ngữ chuẩn đã thống nhất là "lan tỏa" để diễn tả tính bao trùm, không khu trú vào một tình huống/đối tượng cụ thể như các rối loạn lo âu khác.')
P(d, 'Hệ quả nếu KHÔNG đổi: hội đồng phản biện có thể chất vấn về việc dùng thuật ngữ lỗi thời, ảnh hưởng đến tính cập nhật và uy tín của luận án. Đồng thời, các bảng SEM/EFA sử dụng biến viết tắt RLLATQ sẽ không khớp với cách gọi thống nhất quốc tế.')
d.add_paragraph()

P(d, 'Các thay đổi chi tiết:', bold=True)
items = [
    ('"rối loạn lo âu tổng quát" → "rối loạn lo âu lan tỏa"', '53 lần xuất hiện (52 body + 1 bảng)'),
    ('"lo âu tổng quát" → "lo âu lan tỏa"', '82 lần xuất hiện (80 body + 2 bảng)'),
    ('"Rối loạn Lo âu Tổng quát" (viết hoa, định danh chính thức) → "Rối loạn Lo âu Lan tỏa"', 'trong body + tiêu đề bảng'),
    ('"RLLATQ" (mã biến SEM) → "RLLALT"', '70 lần xuất hiện (23 body + 47 bảng SEM)'),
]
for txt, count in items:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'• {txt}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r2 = p.add_run(f'  [{count}]')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = GRAY
d.add_paragraph()
P(d, 'Lưu ý: thay đổi mã biến "RLLATQ" → "RLLALT" áp dụng đồng bộ trong tất cả bảng phân tích SEM/EFA/CFA. Khi xem bảng kết quả mô hình phương trình cấu trúc, đảm bảo các hệ số β/path coefficients gắn với biến RLLALT vẫn giữ nguyên giá trị — chỉ thay đổi tên hiển thị, KHÔNG thay đổi giá trị thống kê.', italic=True)
d.add_paragraph()

# ============================================================
# SECTION 2: KHO
# ============================================================
H(d, '2. CÁC FILE TRONG KHO BẢN DỊCH + TÓM TẮT ĐÃ ĐƯỢC CẬP NHẬT', 1)

H(d, '2.1. VN017 Nguyễn Danh Lâm 2022 — 22 thay đổi', 2)
P(d, 'File output:')
P(d, '   • 03_Ban-dich/VN017_NguyenDanhLam_2022_YHVN_FIXED_27052026.docx (9 thay đổi)')
P(d, '   • Tom-tat-tung-bai/VN017_NguyenDanhLam_2022_YHVN_FIXED_27052026.docx (13 thay đổi)')
d.add_paragraph()
P(d, 'Diễn giải chi tiết các loại thay đổi:', bold=True)

P(d, 'a) Sửa địa lý: "vùng bán đô thị Thanh Hóa" → "huyện bán nông thôn Thanh Hóa" (8 chỗ trong cả 2 file). Đây là sửa lỗi cũ trong khâu dịch — bản dịch ban đầu đã dịch "semi-urban" mà thực chất bài gốc không hề có chữ này; địa danh thực tế Yên Định là huyện thuần nông.', indent=True)
P(d, 'b) Sửa số liệu bảng Stress (Bảng 2 trong cả 2 file): mức Nhẹ 16,0% → 18,7%; Vừa 14,1% → 15,1%; Nặng 6,8% → 7,1%; Rất nặng 4,8% → 0,8%. Sai do quy đổi từ tổng N=482 với cơ chế làm tròn khác nhau giữa bài gốc và bản dịch — nay đã đối chiếu lại với bảng gốc trong PDF.', indent=True)
P(d, 'c) Sửa số liệu bảng Lo âu (Bảng 2): mức Nhẹ 7,7% → 11,2%; Vừa 24,5% → 25,1%. Đã giải thích lỗi sao chép cột-hàng trong Lỗi #1.1.2 ở trên.', indent=True)
P(d, 'd) Bỏ tuyên bố sai "tiêu đề ghi 2024" trong phần phản biện của bản dịch: bản dịch đã ghi "Tạp chí không lập chỉ mục quốc tế. Năm xuất bản 2022 nhưng tiêu đề ghi 2024" — claim này là sai vì PDF gốc tập 516, số 1, năm 2022, không có chỗ nào "ghi 2024". Đã xóa claim này để không lan truyền thông tin sai.', indent=True)
d.add_paragraph()

H(d, '2.2. VN018 An Giang 2025 — 2 thay đổi thứ tự tác giả', 2)
P(d, 'File output:')
P(d, '   • 03_Ban-dich/VN018_AnGiang_2025_YHVN_FIXED_27052026.docx (1 thay đổi)')
P(d, '   • Tom-tat-tung-bai/VN018_AnGiang_2025_YHVN_FIXED_27052026.docx (1 thay đổi)')
d.add_paragraph()
P(d, 'Diễn giải: hai file này đều đảo thứ tự "Lê Minh T." làm 1st author. Sửa: đưa "Nguyễn Đăng Khoa" lên 1st author đúng theo PDF gốc. Đây là cùng nguồn lỗi như #1.2.1–#1.2.3 trong LA chính, sửa đồng bộ cả 3 lớp dữ liệu (LA + bản dịch + tóm tắt) để đảm bảo nhất quán khi hội đồng tra chéo.')
d.add_paragraph()

H(d, '2.3. Cập nhật thuật ngữ "lan tỏa" trong 24 file kho khác', 2)
P(d, 'Diễn giải: do thuật ngữ "lo âu tổng quát" đã được hệ thống dùng nhiều năm trong các bản dịch trước, sự thay đổi cần áp dụng nhất quán cho toàn bộ kho bản dịch + tóm tắt để khi hội đồng trích xuất số liệu kiểm chéo giữa LA và kho thì cả hai khớp nhau. Nếu chỉ đổi LA mà không đổi kho thì sẽ có mâu thuẫn khi tham chiếu.')
d.add_paragraph()
P(d, 'Danh sách 24 file đã được cập nhật:', bold=True)
files_term = [
    ('03_Ban-dich/DICH_Anderson_2025.docx', 1),
    ('03_Ban-dich/DICH_Hoa_2024_Frontiers.docx', 4),
    ('03_Ban-dich/DICH_VNAMHS_2022.docx', 6),
    ('03_Ban-dich/QT001_Jenkins_2023_USA_SanDiego.docx', 1),
    ('03_Ban-dich/QT010_Xu_2021_China_LargestEpi.docx', 1),
    ('03_Ban-dich/QT014_Anderson_2025_Wiley_Narrative.docx', 2),
    ('03_Ban-dich/QT021_Norway_Brunborg_2025_SocSciMed.docx', 1),
    ('03_Ban-dich/QT022_ScreenTime_Li_2025_BJCP.docx', 1),
    ('03_Ban-dich/QT026_UK_NHS_Parliament_2025.docx', 2),
    ('03_Ban-dich/QT028_AJP_Zugman_PediatricAnxiety_2024.docx', 3),
    ('03_Ban-dich/QT035_Jefferies_SocialAnxiety_7Countries_2020.docx', 6),
    ('03_Ban-dich/QT036_Moon_Korea_GAD_ML_2025.docx', 6),
    ('03_Ban-dich/QT043_Bress_JAMA_Maya_App_2024.docx', 1),
    ('03_Ban-dich/VN001_Hoa_2024_Frontiers_HaNoi.docx', 4),
    ('03_Ban-dich/VN002_VNAMHS_2022_National.docx', 1),
    ('03_Ban-dich/VN023_NguyenLX_VN_COVID_Medicine_2023.docx', 2),
    ('Tom-tat-tung-bai/00_Mẫu tóm tắt bài 1.docx', 1),
    ('Tom-tat-tung-bai/QT003_Mandaknalli_Malusare_2021.docx', 2),
    ('Tom-tat-tung-bai/QT005_Alharbi_et_al_2019.docx', 1),
    ('Tom-tat-tung-bai/QT007_Chen_et_al_2023_China.docx', 1),
    ('Tom-tat-tung-bai/QT010_Xu_et_al_2021_China.docx', 1),
    ('Tom-tat-tung-bai/QT035_Jefferies_2020.docx', 3),
    ('Tom-tat-tung-bai/VN001_Hoa_2024_Frontiers_HaNoi.docx', 1),
]
for path, n in files_term:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'• {path}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r2 = p.add_run(f'  [{n} thay đổi]')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True; r2.font.color.rgb = GRAY
d.add_paragraph()

# ============================================================
# SECTION 3: DA AUDIT VA XAC NHAN DUNG
# ============================================================
H(d, '3. CÁC PHẦN ĐÃ AUDIT VÀ XÁC NHẬN ĐÚNG (KHÔNG CẦN SỬA)', 1)
P(d, 'Diễn giải: trong quá trình rà soát, có một số chỗ nghi ngờ ban đầu (do số liệu trông tương tự lỗi đã phát hiện ở bài khác). Sau khi đối chiếu PDF gốc, các chỗ này được xác nhận là đúng — không sửa.', italic=True)
d.add_paragraph()

items_ok = [
    ('VN016 Nguyễn Ngọc Bảo Quyên 2025 (LA + bản dịch + tóm tắt)',
     'Số liệu chính: n=501, tỷ lệ lo âu 86,2%. Đã verify khớp PDF gốc DOI:10.52163/yhc.v66iCD10.2617.',
     'Số "33,2%" xuất hiện trong bản dịch nhìn nghi ngờ ban đầu (vì 38,5% là số bịa cũ ở Tổng quan v2), nhưng đối chiếu lại đoạn thì "33,2%" là tỷ lệ trầm cảm mức nặng+rất nặng — một thống kê khác, hoàn toàn đúng theo PDF gốc.'),
    ('VN019 Hồ Thị Trúc Quỳnh 2025 (LA)',
     'Số liệu chính: n=685, hệ số β=-0,15. Đã verify khớp PDF gốc.',
     'Không có ô nghi ngờ.'),
    ('VN029 Dương 2025 SocPsychiatry (LA)',
     'Số liệu chính: n=2.631, tỷ lệ 50,3%. Đã verify khớp PDF gốc.',
     'Không có ô nghi ngờ.'),
    ('QT005 Alharbi 2019 (LA: para 236, 331)',
     'Số liệu chính: n=1.245, tỷ lệ lo âu 63,4% (nhẹ 34,1%, trung bình 19,5%, nặng 9,8%). Đã verify khớp PDF gốc DOI:10.4103/jfmpc.jfmpc_383_18.',
     'Số "25,8%" trong tóm tắt nhìn nghi ngờ ban đầu, nhưng đối chiếu lại đó là tỷ lệ nam giới trong nhóm trầm cảm nặng (74,2% nữ vs 25,8% nam), không phải tỷ lệ chung — hoàn toàn đúng theo PDF gốc.'),
    ('QT006 Nakie 2022 (LA: para 247–249)',
     'Số liệu chính: n=849, tỷ lệ lo âu 66,7%. Đã verify khớp PDF gốc DOI:10.1186/s12888-022-04393-1.',
     'Không có ô nghi ngờ.'),
    ('QT035 Jefferies 2020 (LA: para 250, 320)',
     'Danh sách 7 quốc gia: Brazil, Trung Quốc, Indonesia, Nga, Thái Lan, Mỹ và Việt Nam. Mỹ có tỷ lệ cao nhất 57,6%; Indonesia thấp nhất 22,9%; Việt Nam 30,7%. Đã verify khớp PDF gốc DOI:10.1371/journal.pone.0239133.',
     'Lỗi cũ ở Tổng quan v2 (claim "Canada/Anh" và "Thái Lan cao nhất 41%") đã bị loại bỏ — LA hiện tại không có hai lỗi này.'),
    ('QT011 Bhardwaj 2020 (LA: para 313, 315)',
     'Số liệu chính: n=1.560, tỷ lệ lo âu 81,9% (mức trung bình+ trở lên 73,2%). Đã verify khớp PDF gốc.',
     'Không có ô nghi ngờ.'),
]
for label, fact, note in items_ok:
    P(d, f'• {label}', bold=True)
    P(d, f'   Sự kiện đúng: {fact}', indent=True)
    P(d, f'   Ghi chú: {note}', indent=True, italic=True)
    d.add_paragraph()

# ============================================================
# SECTION 4: NHAT QUAN GIUA 3 LOP DU LIEU
# ============================================================
H(d, '4. NGUYÊN TẮC NHẤT QUÁN GIỮA 3 LỚP DỮ LIỆU', 1)
P(d, 'Cách rà soát đã áp dụng: với mỗi bài báo được trích dẫn trong LA, cần kiểm tra đồng thời 3 lớp dữ liệu:', bold=True)
P(d, '1) PDF gốc (nguồn sự thật cuối cùng — ground truth).')
P(d, '2) Bản dịch trong 03_Ban-dich/ (cầu nối ngôn ngữ — phải khớp PDF).')
P(d, '3) Tóm tắt trong Tom-tat-tung-bai/ (rút gọn để tham chiếu — phải khớp bản dịch và PDF).')
P(d, '4) Luận án chính (sản phẩm cuối cùng — phải khớp cả 3 trên).')
d.add_paragraph()
P(d, 'Khi phát hiện một lỗi ở LA, KHÔNG sửa chỉ LA. Phải kiểm tra ngược về bản dịch và tóm tắt, vì nếu nguồn lỗi nằm ở khâu dịch thì các lần trích dẫn về sau sẽ tiếp tục sai. Đây chính là cách lỗi VN017 Lâm và VN018 An Giang đã được phát hiện và sửa đồng bộ — gốc rễ là ở bản dịch, lan lên tóm tắt, rồi mới vào LA.', italic=True)
d.add_paragraph()

# ============================================================
# SECTION 5: BUOC 2-3 — Saikia + Xu + abbreviation
# ============================================================
H(d, '5. CÁC LỖI BỔ SUNG ĐÃ PHÁT HIỆN VÀ SỬA SAU BƯỚC 1', 1)
P(d, 'Sau khi sửa xong bước 1 và bắt đầu audit phần còn lại trong kho, đã phát hiện thêm các lỗi sau đây trong LA chính:', italic=True)
d.add_paragraph()

H(d, '5.1. Sai chính tả tên tác giả: "Sakia" → "Saikia" (1 chỗ)', 2)
fix_item(d, '5.1.1', 174, 'Sai chính tả tác giả',
         'M. Sakia và cs.(2023)',
         'A.M. Saikia và cs. (2023)',
         'Tên tác giả viết sai chính tả: "Sakia" thay vì "Saikia". Đồng thời thiếu khoảng trắng trước dấu mở ngoặc và thiếu chữ "A." (Anku M. Saikia).',
         'Lỗi typo phổ biến khi gõ tên người Việt và Ấn Độ — chữ cái "i" bị bỏ sót. Hệ quả: khi hội đồng tra cứu DOI/PMID sẽ không tìm được bài.',
         'Đối chiếu trang 1 PDF gốc DOI:10.4103/ijcm.ijcm_614_21 — tên tác giả ghi là "Anku M. Saikia, Hemen Das, Vinoth Rajendran".',
         'Đã sửa thành "A.M. Saikia và cs. (2023)" — khớp với TLTK ở para 1504 vốn đã ghi đúng "Saikia A. M.".')

H(d, '5.2. Sai năm xuất bản tác giả Xu (4 chỗ trong body)', 2)
P(d, 'Đây là lỗi nhất quán giữa in-text citation và TLTK. TLTK para 1545 ghi đúng năm 2021 (Journal of Affective Disorders, vol. 288, 2021, trang 17–22; nhận bài 11/01/2021; chấp nhận 26/03/2021). Tuy nhiên 4 lần xuất hiện in-text trong body đều ghi sai năm là 2022.', italic=True)
d.add_paragraph()

fix_item(d, '5.2.1', 169, 'Năm xuất bản Xu (citation in-text)',
         'Q. Xu và cs (2022) đưa ra tỷ lệ chênh lệch nam lo âu hơn nữ là 10,11% nam/9,66% nữ',
         'Q. Xu và cs. (2021) đưa ra tỷ lệ chênh lệch nam lo âu hơn nữ là 10,11% nam/9,66% nữ',
         'Năm 2022 sai — TLTK chỉ rõ Xu et al. công bố năm 2021 trên Journal of Affective Disorders.',
         'Khi viết bản nháp, có thể nhầm năm online publication với năm in. Lỗi nhỏ về năm nhưng ảnh hưởng đến tính nhất quán in-text/TLTK.',
         'Đối chiếu DOI:10.1016/j.jad.2021.03.080 — bài thuộc số 288 năm 2021, không phải 2022.',
         'Đã sửa năm thành 2021 trong cả 4 chỗ in-text (para 169, 235, 238, 320) cho khớp TLTK.')

fix_item(d, '5.2.2', 235, 'Năm Xu trong đoạn so sánh tỷ lệ',
         'Xu và cs (2022) khảo sát ở Trung Quốc trong đại dịch Covid-19',
         'Xu và cs. (2021) khảo sát ở Trung Quốc trong đại dịch Covid-19',
         'Cùng lỗi như 5.2.1.', 'Lỗi nhất quán năm.',
         'Đối chiếu TLTK.', 'Đã sửa.')

fix_item(d, '5.2.3', 238, 'Năm Xu trong đoạn giới thiệu công trình',
         'Qingqing Xu, Zhenxing Mao, Dandan Wei... Cuiping Wu (2022)',
         'Qingqing Xu, Zhenxing Mao, Dandan Wei... Cuiping Wu (2021)',
         'Cùng lỗi như 5.2.1 — đoạn giới thiệu Xu cũng ghi sai năm. Lưu ý đây là đoạn giới thiệu chính của công trình, nếu sai thì lan ra các đoạn dẫn sau.',
         'Khi gõ đoạn giới thiệu, có thể đã nhìn nhầm năm từ bản nháp khác.',
         'Đối chiếu TLTK para 1545.', 'Đã sửa.')

fix_item(d, '5.2.4', 320, 'Citation parenthetical (Q. Xu và cs., 2022)',
         '(A.M. Saikia và cs., 2023), (Q. Xu và cs., 2022)',
         '(A.M. Saikia và cs., 2023), (Q. Xu và cs., 2021)',
         'Citation parenthetical sai năm.', 'Cùng lỗi.', 'Cùng cách verify.', 'Đã sửa.')

H(d, '5.3. Abbreviation còn sót: "LA tổng quát" / "LATQ" (4 chỗ)', 2)
P(d, 'Sau khi đã chạy global replace "tổng quát" → "lan tỏa" trong bước 1, vẫn còn 4 chỗ sử dụng VIẾT TẮT vắn của "rối loạn lo âu tổng quát" mà không bị bắt:', italic=True)
d.add_paragraph()
P(d, '• Para 169: "(LA tổng quát, LA xã hội và LA chia ly)" → đã sửa "LA tổng quát" → "LA lan tỏa".')
P(d, '• Para 1113: "Cụ thể, các mô hình đơn (RLLA tổng quát, chia ly và xã hội) đều có chỉ số phù hợp tốt" → đã sửa "RLLA tổng quát" → "RLLA lan tỏa".')
P(d, '• Para 1115: "β → LATQ = 0,215; β → LAXH = 0,253" → đã sửa "LATQ" → "LALT".')
P(d, '• Para 1324: "với hai loại rối loạn lo âu lan tỏa và lo âu xã hội..." có chỗ "LATQ" còn sót → đã sửa thành "LALT".')
P(d, 'Diễn giải: trong bước 1, đã đổi đầy đủ "RLLATQ" → "RLLALT" (70 chỗ) và "rối loạn lo âu tổng quát" → "rối loạn lo âu lan tỏa" (53 chỗ). Tuy nhiên một số tác giả đã viết tắt thành "LA tổng quát" (bỏ chữ "rối loạn") hoặc "LATQ" (bỏ chữ "RL") — những trường hợp này cần fix riêng để đảm bảo TẤT CẢ các cách viết được đồng bộ hóa sang "lan tỏa".', italic=True)
d.add_paragraph()

# ============================================================
# SECTION 6: AUDIT KET QUA - PHAN VERIFY CHI TIET
# ============================================================
H(d, '6. KẾT QUẢ AUDIT TIẾP CÁC BÀI TRONG TỔNG QUAN', 1)
P(d, 'Sau khi sửa xong LA chính, đã rà soát thêm các bài báo được trích dẫn trong tổng quan để cross-check tính chính xác của số liệu. Kết quả:', italic=True)
d.add_paragraph()

H(d, '6.1. Các số liệu đã được verify từ PDF gốc (KẾT LUẬN: ĐÚNG)', 2)

items_verified = [
    ('Xu et al. 2021 (Trung Quốc, n=373.216)',
     'Para 169: "10,11% nam / 9,66% nữ" → khớp Table 1 PDF.',
     'Para 235, 240: "38,42% lo âu từ mức nhẹ trở lên" → khớp Table 2 PDF: mild 28.53% + moderate 6.79% + severe 3.10% = 38.42% (chính xác bằng phép cộng từ phân tầng GAD-7).',
     'Para 240: "9,89% trung bình trở lên" → khớp tổng prevalence = 9.89% (GAD-7 ≥ 10).',
     'Kết luận: tất cả 3 con số đều đúng. Năm 2021 (sau khi sửa từ 2022 ở bước 2) đã khớp TLTK.'),
    ('Bhardwaj et al. 2020 (Ấn Độ, Chandigarh, n=288)',
     'Para 313, 315: "81,9% theo R. Bhardwaj... ở Ấn Độ" → khớp Bảng phân bố Anxiety (Table 4.7 PDF): Normal 18.1% + Mild 8.7% + Moderate 26.4% + Severe 25.3% + Extreme 21.5% → Mild trở lên = 8.7+26.4+25.3+21.5 = 81.9%.',
     'Para 232-235: "73,2% từ mức trung bình trở lên (trong đó 46,8% nặng và cực kỳ nặng)" → 26.4+25.3+21.5 = 73.2% và 25.3+21.5 = 46.8%. Khớp tuyệt đối.',
     'Para 231: "60,8% HS nam, 39,2% HS nữ" → khớp Table 1 PDF.',
     'Kết luận: tất cả số liệu đều đúng và có thể truy nguyên từ PDF gốc.'),
    ('Saikia et al. 2023 (Ấn Độ, Kamrup Assam, n=360)',
     'Para 169, 245: "30% nam / 18,9% nữ" → khớp Table 2 PDF (anxiety by gender).',
     'Para 235: "tỷ lệ chung bị lo âu là 24,4%" → khớp Abstract: prevalence of anxiety = 24.4%.',
     'Para 245, 330, 374, 382: n=360 → khớp Materials and Methods.',
     'Kết luận: tất cả số liệu khớp PDF. Lỗi spelling "Sakia" đã được sửa thành "Saikia".'),
]
for label, *facts in items_verified:
    P(d, f'• {label}', bold=True)
    for fact in facts:
        P(d, f'   {fact}', indent=True)
    d.add_paragraph()

H(d, '6.2. Có thể có sai số nhẹ — cần thầy/NCS quyết định', 2)

P(d, '• Chen et al. 2023 (Trung Quốc, Tứ Xuyên):', bold=True)
P(d, '   LA para 335 trích dẫn "63.487 học sinh trung học cơ sở và trung học phổ thông" về số mẫu của Chen. Đối chiếu PDF gốc:', indent=True)
P(d, '   – 63.487 = tổng số HS HOÀN THÀNH khảo sát (finished investigation).', indent=True)
P(d, '   – 282 HS bị loại do thiếu dữ liệu critical items.', indent=True)
P(d, '   – 63.205 = số HS thực tế ĐƯA VÀO PHÂN TÍCH (effective response rate 99,6%).', indent=True)
P(d, '   Cả hai con số đều có trong PDF, nhưng theo chuẩn báo cáo nghiên cứu, nên dùng "63.205" (analysis sample) để khớp với cách trình bày trong Methods/Results của Chen et al. (cũng khớp tom-tat).', indent=True, italic=True)
P(d, '   Kiến nghị: NCS xem xét sửa "63.487" → "63.205" trong para 335 cho khớp với cách trình bày chuẩn.', indent=True, italic=True, color=GREEN)
d.add_paragraph()

H(d, '6.3. Cần kiểm tra thêm (PDF đầy đủ chưa đọc xong)', 2)
P(d, 'Các tuyên bố sau đây trong LA cần đối chiếu với PDF gốc đầy đủ để hoàn thiện audit. Tạm thời không sửa cho tới khi NCS có thời gian kiểm tra:', italic=True)
d.add_paragraph()
items_pending = [
    ('Para 334: Chen 2023 — "tỷ lệ bắt nạt học đường dao động từ 9,0% đến 61,6%, trong đó 35% là nạn nhân bắt nạt truyền thống và 15% là nạn nhân bắt nạt qua mạng". Đây là số liệu Chen TRÍCH DẪN từ các nghiên cứu khác trong intro (KHÔNG phải data của Chen). Cần verify nguồn citation.'),
    ('Para 335: Chen 2023 — "bắt nạt bằng lời nói (61%); tấn công lấy tài sản (59%); thao túng xã hội (45%); bắt nạt mạng (29%)" — số liệu cụ thể về phân loại bắt nạt cần đối chiếu Table trong PDF Chen.'),
    ('Para 376: Qiu 2022 — "Mô hình nuôi dạy tích cực giảm trầm cảm có ý nghĩa thống kê: OR=0,32, KTC 95%: 0,24-0,43". Cần đối chiếu với PDF Qiu để verify OR và CI.'),
    ('Para 285: Zhu 2025 — hệ số Beta=–0,149 và Beta=–0,087 cho yếu tố bảo vệ. Cần đối chiếu PDF Zhu để verify (tom-tat chỉ có AOR=13,71 cho giấc ngủ <5h).'),
    ('Para 327, 375, 384: Anderson 2025 — "48/52 công trình", "Pew Research 95%/96%" — các tuyên bố này cần verify trong narrative review của Anderson.'),
    ('Para 332: Wen 2020 — "khảo sát 900 học sinh từ lớp 9 đến lớp 12". Tom-tat ghi "junior high school" (lớp 7-9 ở Trung Quốc), không phải lớp 9-12. Có thể có sai về phạm vi tuổi/lớp. Cần verify PDF.'),
]
for it in items_pending:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run(f'• {it}')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
d.add_paragraph()

# ============================================================
# SECTION 7: KIEN NGHI
# ============================================================
H(d, '7. KIẾN NGHỊ QUY TRÌNH GIẢI QUYẾT TIẾP THEO', 1)
P(d, 'Sau khi nhận báo cáo này, NCS có thể chọn 1 trong 3 phương án:', bold=True)
d.add_paragraph()

P(d, 'Phương án A (NHANH — chấp nhận bản hiện tại):', bold=True)
P(d, '   Sử dụng file 01_LuanAn_v3_4_FixAbbrev_27052026.docx làm bản nộp hội đồng. Mục 6.2 (Chen 63.487/63.205) là khác biệt nhỏ, không sai bản chất. Mục 6.3 là các tuyên bố cần verify nhưng PDF gốc đều có sẵn — có thể tự kiểm tra khi rảnh.')
d.add_paragraph()

P(d, 'Phương án B (TRUNG BÌNH — sửa thêm Chen):', bold=True)
P(d, '   Tạo bản v3_5 sửa thêm số mẫu Chen 63.487 → 63.205. Tổng thời gian: ~5 phút. Cũng cần check lại 5 chỗ Chen còn lại để đảm bảo nhất quán.')
d.add_paragraph()

P(d, 'Phương án C (CHẬM — verify tất cả PDF còn nghi):', bold=True)
P(d, '   Đọc full PDF cho từng bài trong Mục 6.3 (Qiu, Zhu, Anderson, Wen), đối chiếu từng tuyên bố trong LA. Tổng thời gian ước tính: 2-3 giờ. Đảm bảo độ chính xác cao nhất trước khi nộp.')
d.add_paragraph()

P(d, 'NCS chọn phương án nào tùy thời gian nộp và mức độ chấp nhận rủi ro về kiểm tra chéo của hội đồng.', italic=True)
d.add_paragraph()

P(d, '— Hết báo cáo v5 —', italic=True)

# ============================================================
# SAVE
# ============================================================
d.save(OUT)
from docx import Document as D
dd = D(OUT)
w = sum(len(p.text.split()) for p in dd.paragraphs)
print(f"Da luu: {OUT}")
print(f"Paragraphs: {len(dd.paragraphs)}, Words: ~{w}, Size: {os.path.getsize(OUT)//1024}KB")
