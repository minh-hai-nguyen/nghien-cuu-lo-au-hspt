"""DOC: Tra loi cau hoi 'F-test ANOVA gioi la gi?' don gian, de hieu cho thay.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/F_test_ANOVA_gioi_la_gi.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('F-TEST ANOVA GIỚI LÀ GÌ?\n— Trả lời câu hỏi của thầy 09/05/2026 —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue('F-test ANOVA giới là gì?', italic=True)

# Trả lời ngắn
H('Trả lời NGẮN GỌN (1 câu)', level=2, color=NAVY)
para(
    'F-test ANOVA giới là kiểm định thống kê dùng để TRẢ LỜI câu '
    'hỏi: "Mức độ lo âu CÓ khác biệt có ý nghĩa thống kê giữa NAM '
    'và NỮ hay không?"', bold=True
)

# Phân tích từng thành phần
H('Phân tích từng thành phần của thuật ngữ', level=2, color=NAVY)
add_table(
    ['Thành phần', 'Tiếng Anh', 'Ý nghĩa đơn giản'],
    [
        ['ANOVA', 'Analysis of Variance',
         'Phân tích phương sai — kỹ thuật so sánh ĐIỂM TRUNG BÌNH giữa các nhóm'],
        ['F-test', 'F-test (Fisher\'s test)',
         'Phép kiểm định thống kê dùng trong ANOVA — cho ra một con số gọi là "F"'],
        ['Giới', 'Gender / Sex',
         'Biến độc lập trong phân tích này — có 2 nhóm: nam và nữ'],
    ]
)
para('')
para(
    'GHÉP LẠI: "F-test ANOVA giới" = dùng F-test trong ANOVA để '
    'kiểm tra xem điểm trung bình lo âu của NAM và NỮ có KHÁC '
    'NHAU đáng kể không.', bold=True
)

# Cách hoạt động
H('Cách F-test hoạt động (đơn giản hóa)', level=2, color=NAVY)
para('F-test tính tỷ số:', bold=True)
para('     F = (Phương sai GIỮA các nhóm) / (Phương sai TRONG các nhóm)')
para('Diễn giải đơn giản:')
para('• Tử số: ĐTB nam và ĐTB nữ KHÁC NHAU bao nhiêu?')
para('• Mẫu số: TRONG nhóm nam (và nhóm nữ), điểm các học sinh phân tán bao nhiêu?')
para('• F càng LỚN → khác biệt giữa nam-nữ càng RÕ so với phân tán nội nhóm.')
para('')
para('Quy tắc đọc F-test:', bold=True)
add_table(
    ['Giá trị F', 'p-value', 'Kết luận'],
    [
        ['F lớn (vd > 4)', 'p < 0,05', 'CÓ khác biệt có ý nghĩa thống kê giữa nam-nữ'],
        ['F nhỏ (vd < 1)', 'p > 0,05', 'KHÔNG có khác biệt có ý nghĩa thống kê'],
    ]
)

# Áp dụng vào Bảng 3.20
H('Áp dụng vào Bảng 3.20 chương 3 luận án', level=2, color=NAVY)
para(
    'Trong Bảng 3.20, tác giả chạy F-test ANOVA giới CHO TỪNG '
    'DẠNG RỐI LOẠN LO ÂU (RLLATQ, RLLACL, RLLAXH) — kết quả:'
)
add_table(
    ['Dạng RLLA', 'Nam (n=614) ĐTB', 'Nữ (n=738) ĐTB', 'F-test', 'p-value', 'Kết luận'],
    [
        ['Lo âu lan tỏa (RLLATQ)', '51,43', '59,47', '44,484', '< 0,001',
         'CÓ khác biệt — nữ > nam RÕ RỆT'],
        ['Lo âu chia ly (RLLAC)', '25,42', '24,76', '0,246', '0,620',
         'KHÔNG khác biệt — nam ≈ nữ'],
        ['Lo âu xã hội (RLLAXH)', '43,20', '52,74', '45,984', '< 0,001',
         'CÓ khác biệt — nữ > nam RÕ RỆT'],
    ]
)
para('')
para('Diễn giải từng kết quả:', bold=True)
para('• F (giới × RLLATQ) = 44,484 — RẤT LỚN, p < 0,001 → khác biệt nữ-nam ở lo âu lan tỏa CỰC KỲ có ý nghĩa.')
para('• F (giới × RLLAXH) = 45,984 — RẤT LỚN, p < 0,001 → khác biệt nữ-nam ở lo âu xã hội CỰC KỲ có ý nghĩa.')
para('• F (giới × RLLAC) = 0,246 — RẤT NHỎ, p = 0,620 (> 0,05) → KHÔNG có khác biệt nữ-nam ở lo âu chia ly.')

# Lưu ý quan trọng
H('Lưu ý quan trọng — F-test KHÁC effect size', level=2, color=NAVY)
para(
    'F-test trả lời câu hỏi: "CÓ khác biệt KHÔNG?" (yes/no theo '
    'p-value).', bold=True
)
para(
    'F-test KHÔNG trả lời: "Khác biệt LỚN NHƯ THẾ NÀO?" — câu này '
    'phải dùng Cohen d.', bold=True
)
para('Ví dụ với chương 3 luận án:')
add_table(
    ['Dạng RLLA', 'F-test', 'Cohen d (effect size chuẩn)', 'Diễn giải'],
    [
        ['RLLATQ', '44,484', '0,365', 'Có khác biệt; cường độ small-to-medium'],
        ['RLLAXH', '45,984', '0,370', 'Có khác biệt; cường độ small-to-medium'],
    ]
)
para('')
para(
    'Tuy F-test RLLAXH (45,984) > F-test RLLATQ (44,484), Cohen d '
    'gần như BẰNG NHAU (0,370 vs 0,365). Lý do: F-test phụ thuộc '
    'cả PHƯƠNG SAI trong nhóm — RLLAXH có SD lớn hơn, làm F tăng '
    'nhưng không có nghĩa effect size lớn hơn. Khi báo cáo cần '
    'NÊU CẢ HAI: F-test (xem có khác biệt không) + Cohen d '
    '(cường độ khác biệt).', italic=True
)

# Tóm gọn
H('TÓM GỌN', level=2, color=NAVY)
blue('F-test ANOVA giới = công cụ thống kê trả lời câu hỏi: "Lo âu '
     'có khác biệt giữa nam và nữ không?"', bold=True)
blue('Kết quả Bảng 3.20:', bold=True)
blue('• Lo âu lan tỏa: F = 44,484 (p < 0,001) → CÓ khác biệt, nữ > nam')
blue('• Lo âu chia ly: F = 0,246 (p = 0,620) → KHÔNG có khác biệt')
blue('• Lo âu xã hội: F = 45,984 (p < 0,001) → CÓ khác biệt, nữ > nam')
blue('Cảnh báo: F-test KHÔNG đo cường độ — phải kết hợp với Cohen '
     'd. Cohen d cho RLLATQ (0,365) và RLLAXH (0,370) gần như bằng '
     'nhau — chênh lệch giới ở hai dạng TƯƠNG ĐƯƠNG về cường độ '
     'thực tế.')

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
