# -*- coding: utf-8 -*-
"""Part 3: MENTAL HEALTH chapter (pages 17-24) — full translation."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_imgs/'

doc = Document(OUT)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def H2(text):
    return P(text, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(text):
    return P(text, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))
def H4(text):
    return P(text, bold=True, italic=True, size=11)

def PageMarker(page_num, note='V-NAMHS, UNICEF/IOS 2022'):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'--- Trang {page_num}, {note} ---')
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

def AddImg(filename, caption_vn, caption_en, width_cm=11.0):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(path, width=Cm(width_cm))
    except Exception as e:
        print(f'  [ERR] {filename}: {e}')
        return
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    hr = t.rows[0]
    for i, h in enumerate(headers):
        c = hr.cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# =============================================================
# MENTAL HEALTH chapter (pages 17-24)
# =============================================================
PageMarker('17')
H2('SỨC KHOẺ TÂM THẦN (MENTAL HEALTH)')

H3('Tổng quan')
P('Vị thành niên 10–19 tuổi chiếm 14,3 % dân số Việt Nam (Tổng cục Thống kê 2020). Tuy nhiên, rất ít biết về tỷ lệ rối loạn tâm thần ở VTN Việt Nam. Đây là khoảng trống quan trọng vì nghiên cứu quốc tế cho thấy rối loạn tâm thần trong giai đoạn vị thành niên — đặc biệt khi không được điều trị hoặc điều trị chưa đủ — có thể kéo theo hậu quả tiêu cực suốt đời (Erskine et al. 2016, Ormel 2017). Một rà soát tài liệu trước đó tìm thấy các ước tính tỷ lệ vấn đề SKTT ở VN dao động rộng từ 8 % đến 29 % (Samuels et al. 2018).')

P('Tuy nhiên, các nghiên cứu sẵn có về rối loạn tâm thần VTN tại VN thường bị hạn chế bởi cỡ mẫu nhỏ, địa điểm hoặc phạm vi tuổi bị hạn chế, đo một hoặc một vài rối loạn, và thiếu nghiên cứu dùng công cụ chẩn đoán (instruments được thiết kế để xác định rối loạn tâm thần theo tiêu chí đã thiết lập). Điều này khiến khả năng khái quát và ảnh hưởng chính sách của các nghiên cứu này bị hạn chế. Ví dụ, nghiên cứu đại diện quốc gia về SKTT ở trẻ 6–16 tuổi của Weiss và cộng sự (2014) dùng SDQ đo 1.314 phụ huynh và 591 vị thành niên (trong mẫu 12–16 tuổi). Trong mẫu vị thành niên cụ thể, 10,7 % (tự báo cáo) đến 11,9 % (phụ huynh báo cáo) đạt ngưỡng "borderline" cho vấn đề SKTT. Tỷ lệ tương tự về vấn đề cảm xúc và hành vi (12,4 % và 11,3 %) được tìm thấy trong cùng nghiên cứu khi dùng YSR và CBCL. Tuy nhiên, các công cụ này là thang sàng lọc triệu chứng, không phải công cụ chẩn đoán, do đó không đo được tỷ lệ rối loạn tâm thần theo tiêu chí chẩn đoán.')

P('Ví dụ khác, Samuels và cộng sự (2018) dùng SDQ đo SKTT ở VTN 12–17 tuổi sống tại các thành phố (Hà Nội, TP HCM, Điện Biên Phủ, Long Xuyên) và vùng nông thôn (Kéo Lôm và Phú Mỹ). Trong 402 VTN được khảo sát, 13,4 % đạt ngưỡng "bất thường" của SDQ. Mẫu tương đối nhỏ (n = 402) khiến khó khái quát lên quần thể VTN VN. Mặc dù các nghiên cứu này đóng góp hiểu biết quan trọng, chúng không thể khái quát để thông tin tỷ lệ rối loạn tâm thần ở cấp quốc gia. Do đó, bằng chứng cho hoạch định chính sách SKTT VTN VN và khả năng hành động hiệu quả là rất hạn chế.')

P('Mục tiêu cốt lõi của V-NAMHS là tạo ra ước tính đại diện quốc gia về tỷ lệ rối loạn tâm thần ở VTN VN. Chương này trình bày đo lường rối loạn tâm thần trong V-NAMHS và các phát hiện. Báo cáo trình bày cả vấn đề SKTT và rối loạn tâm thần; ngoài ra, tỷ lệ hành vi tự sát và tự làm hại bản thân cũng được trình bày trong bối cảnh SKTT. Tất cả tỷ lệ và số lượng trong chương này đều được gia trọng.')

PageMarker('18')
H3('Đo lường (Measurement)')

H4('Các công cụ (Measures)')
P('Các module chẩn đoán của DISC-5 được dùng để đánh giá tỷ lệ 12 tháng qua của các rối loạn được chọn. DISC-5 là công cụ chẩn đoán chuẩn hoá (Bitsko et al. 2019, Shaffer et al. 2000) do Đại học Columbia phát triển ban đầu với hỗ trợ của Viện Sức khoẻ Tâm thần Quốc gia Hoa Kỳ (NIMH). DISC-5 được thiết kế để phỏng vấn viên "lay" (không có đào tạo lâm sàng nhưng được huấn luyện DISC-5) thực hiện. DISC-5 được dùng để đo 6 rối loạn tâm thần trong V-NAMHS: ám ảnh xã hội, rối loạn lo âu lan toả, rối loạn trầm cảm nặng, PTSD, rối loạn hành vi và ADHD. Các rối loạn này được chọn vì phổ biến trong giai đoạn vị thành niên và chịu trách nhiệm cho tỷ lệ gánh nặng rối loạn tâm thần đáng kể ở nhóm tuổi này (Erskine et al. 2015). Ám ảnh xã hội và rối loạn lo âu lan toả được gộp chung thành "các rối loạn lo âu" (anxiety disorders) trong báo cáo này. Ngoại trừ ADHD (hỏi phụ huynh), tất cả các module DISC-5 khác được hỏi trực tiếp vị thành niên.')

P('Các công cụ đo hành vi tự sát và tự làm hại bản thân cũng được bao gồm. Hành vi tự sát đề cập đến: (1) ý nghĩ tự sát (suy nghĩ nghiêm túc về việc kết thúc cuộc sống của mình); (2) lập kế hoạch tự sát (lên kế hoạch để kết thúc cuộc sống); (3) toan tự sát (khi tự gây hại với ý định kết thúc cuộc sống). Tự làm hại bản thân (self-harm) nghĩa là cố ý làm đau hay gây tổn thương cho bản thân mà không có ý định kết thúc cuộc sống (tức non-suicidal self-injury). Tất cả vị thành niên được hỏi về hành vi tự sát và tự làm hại bất kể tình trạng rối loạn tâm thần. Đối với hành vi tự sát trong 12 tháng qua, VTN chỉ được hỏi về lập kế hoạch nếu đã báo cáo ý nghĩ tự sát, và chỉ được hỏi về toan tự sát nếu đã báo cáo lập kế hoạch.')

H4('Vấn đề SKTT và rối loạn tâm thần (Mental health problems and mental disorders)')
P('Nội dung và cấu trúc DISC-5 được thiết kế tuân theo tiêu chí chẩn đoán đã thiết lập cho rối loạn tâm thần và được cập nhật gần đây để đáp ứng Sổ tay Chẩn đoán và Thống kê Rối loạn Tâm thần, Tái bản thứ 5 (DSM-5) (American Psychiatric Association 2013). Chẩn đoán rối loạn tâm thần theo DSM-5 yêu cầu cá nhân đáp ứng các triệu chứng cốt lõi của một rối loạn tâm thần nhất định và đạt ngưỡng cụ thể về thời gian, tần suất, mức độ nghiêm trọng và biểu hiện của các triệu chứng, đồng thời báo cáo mức độ suy giảm chức năng tối thiểu do các triệu chứng này gây ra. Các câu hỏi trong DISC-5 phản ánh điều này bằng cách đầu tiên hỏi về một triệu chứng nhất định rồi hỏi tiếp các câu hỏi chi tiết về đặc điểm cụ thể của triệu chứng đó theo yêu cầu chẩn đoán.')

PageMarker('19')
P('Điều này phân biệt một công cụ chẩn đoán, như DISC-5, với các "thang triệu chứng" (symptom scales) chỉ hỏi về sự hiện diện chung của triệu chứng. Vì vậy, thang triệu chứng thường báo cáo tỷ lệ cao hơn nhiều so với công cụ chẩn đoán (COVID-19 Mental Disorders Collaborators 2021, Ferrari et al. 2013, GBD 2019 Mental Disorders Collaborators 2022) vì chúng đánh giá triệu chứng, không phải rối loạn.')

P('Tuy nhiên, các cá nhân vẫn có thể trải qua căng thẳng và suy giảm chức năng liên quan mà không nhất thiết đáp ứng tiêu chí DSM-5 để chẩn đoán một rối loạn cụ thể. Những cá nhân này có thể là nhóm quan trọng cho can thiệp trước khi tiến triển thành một rối loạn tâm thần đầy đủ (Pagliaro et al. 2021). Song song, đã có câu hỏi về tính áp dụng của tiêu chí chẩn đoán DSM-5 với các nền văn hoá phi phương Tây cũng như tác động của các yếu tố văn hoá khi sử dụng công cụ chuẩn hoá như DISC-5 (Canino and Alegria 2008).')

P('Thừa nhận các thách thức tiềm tàng này, hai bộ phát hiện được trình bày. Thứ nhất, tỷ lệ "vấn đề SKTT" (mental health problems) được báo cáo, bao gồm các cá nhân đáp ứng ít nhất một nửa số triệu chứng cho một rối loạn tâm thần theo DISC-5 nhưng có thể không nhất thiết đáp ứng tất cả tiêu chí chẩn đoán DSM-5. Thứ hai, tỷ lệ "rối loạn tâm thần" (mental disorders) được báo cáo. Nhóm này gồm các cá nhân đáp ứng tiêu chí chẩn đoán DSM-5. Bảng 5 dưới đây đưa ra định nghĩa chung cho vấn đề SKTT và rối loạn tâm thần, cũng như định nghĩa thao tác hóa áp dụng trong V-NAMHS. Thuật ngữ cho các loại vấn đề SKTT khác với rối loạn tâm thần cũng được hiển thị.')

MakeTable(
    ['', 'Vấn đề SKTT (Mental health problem)', 'Rối loạn tâm thần (Mental disorder)'],
    [
        ('Định nghĩa chung',
         'Can thiệp vào cách một người suy nghĩ, cảm xúc và hành vi, nhưng ở mức độ ít hơn so với rối loạn tâm thần. Có thể là tạm thời hoặc phản ứng cấp tính với căng thẳng cuộc sống.',
         'Hội chứng hành vi hoặc tâm lý có ý nghĩa lâm sàng xảy ra ở một cá nhân, gắn với căng thẳng hiện tại (triệu chứng đau đớn), khuyết tật (suy giảm ở một hoặc nhiều lĩnh vực chức năng quan trọng), và/hoặc tăng nguy cơ tử vong, đau đớn, khuyết tật hoặc mất tự do đáng kể.'),
        ('Định nghĩa áp dụng trong V-NAMHS',
         'Một VTN được xem là có vấn đề SKTT nếu đáp ứng ít nhất một nửa số triệu chứng cần thiết để chẩn đoán một rối loạn tâm thần (tức "subthreshold symptoms"). Do đó, nhóm có vấn đề SKTT bao gồm các em có triệu chứng dưới ngưỡng hoặc đủ ngưỡng, có hoặc không suy giảm chức năng. Tức, nhóm này cũng bao gồm các em đáp ứng tiêu chí chẩn đoán rối loạn.',
         'Một VTN được xem là đáp ứng tiêu chí DSM-5 nếu tất cả các triệu chứng yêu cầu (tức "full threshold symptoms") và mức độ suy giảm chức năng do các triệu chứng này đều được ghi nhận. Chẩn đoán này tuân theo các thuật toán chấm điểm chuẩn do tác giả DISC-5 cung cấp.'),
        ('Thuật ngữ cho các loại khác nhau',
         'Trầm cảm (Depression) — Lo âu (Anxiety) — Căng thẳng sau sang chấn (Posttraumatic stress) — Các vấn đề hành vi (Conduct problems) — Các vấn đề mất chú ý và/hoặc tăng động (Problems with inattention and/or hyperactivity)',
         'Rối loạn trầm cảm nặng (Major depressive disorder) — Các rối loạn lo âu (Anxiety disorders) — PTSD — Rối loạn hành vi (Conduct disorder) — ADHD'),
    ])
P('Bảng 5. Định nghĩa vấn đề SKTT và rối loạn tâm thần. Nguồn: Bảng 5 báo cáo gốc, trang 19.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

PageMarker('19-end')

# image on p19
AddImg('p019_img1_Image216.jpg',
       'Hình minh hoạ Chương Sức khoẻ tâm thần',
       'Chapter divider — Mental Health section',
       width_cm=11.0)

PageMarker('20')
H3('Phát hiện (Findings)')
H4('Vấn đề SKTT (Mental health problems)')
P('Một phần năm vị thành niên (21,7 %) báo cáo có một vấn đề SKTT trong 12 tháng qua (Bảng 6). Không có khác biệt về tỷ lệ giữa nam và nữ, hoặc giữa nhóm 10–13 tuổi và 14–17 tuổi.')

MakeTable(
    ['Vấn đề SKTT', '% 10–13', 'n 10–13', '% 14–17', 'n 14–17', '% 10–17', 'n 10–17'],
    [
        ('Nam', '21,4', '357', '20,2', '292', '20,8', '649'),
        ('Nữ', '20,2', '302', '25,3', '350', '22,6', '651'),
        ('Tổng', '20,8', '659', '22,7', '462', '21,7', '1.301'),
    ])
P('Bảng 6. Tỷ lệ 12 tháng qua của vấn đề SKTT ở VTN 10–17 tuổi theo giới và nhóm tuổi. Gia trọng N: nam = 3.119; nữ = 2.877; 10–13 = 3.167; 14–17 = 2.829. Nguồn: Bảng 6 báo cáo gốc, trang 20.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('Hình 1 cho thấy lo âu có tỷ lệ cao nhất trong các vấn đề SKTT (18,6 %).')

# Figure 1 text summary (image not always extractable as chart)
MakeTable(
    ['Loại vấn đề SKTT', 'Tỷ lệ %'],
    [
        ('Lo âu (Anxiety)', '18,6'),
        ('Trầm cảm (Depression)', '4,3'),
        ('Các vấn đề mất chú ý và/hoặc tăng động (Problems with inattention/hyperactivity)', '2,8'),
        ('Căng thẳng sau sang chấn (Posttraumatic stress)', '1,0'),
        ('Các vấn đề hành vi (Conduct problems)', '0,7'),
    ],
    header_bg='FFF2CC')
P('Hình 1. Tỷ lệ 12 tháng qua của vấn đề SKTT ở VTN 10–17 tuổi theo loại. Nguồn: Figure 1 báo cáo gốc, trang 20.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('VTN có vấn đề SKTT sau đó được phân tích thêm để xác định tỷ lệ báo cáo suy giảm chức năng do triệu chứng, kết hợp với tỷ lệ báo cáo triệu chứng ở ngưỡng đầy đủ hay dưới ngưỡng. Ngưỡng triệu chứng dựa trên tiêu chí DSM-5, trong đó triệu chứng dưới ngưỡng (subthreshold) cho biết đáp ứng ít nhất một nửa số triệu chứng yêu cầu, trong khi triệu chứng đầy đủ ngưỡng (full threshold) cho biết đáp ứng tất cả triệu chứng yêu cầu. Điều này có nghĩa VTN có vấn đề SKTT được phân thành một trong bốn nhóm trong Hình 2 (lưu ý: triệu chứng đầy đủ ngưỡng CÓ suy giảm = đáp ứng tiêu chí rối loạn tâm thần). Như Hình 2 cho thấy, hầu hết VTN có vấn đề SKTT báo cáo có mức suy giảm nào đó do triệu chứng của họ — cho dù ghi nhận tất cả triệu chứng (3,3 %) hay ít nhất một nửa (10,3 %).')

PageMarker('21')
MakeTable(
    ['Phân nhóm vấn đề SKTT', 'Tỷ lệ %'],
    [
        ('Không đạt ngưỡng triệu chứng (No symptom threshold met)', '78,3'),
        ('Triệu chứng đầy đủ + suy giảm = RỐI LOẠN TÂM THẦN', '3,3'),
        ('Triệu chứng dưới ngưỡng + suy giảm', '10,3'),
        ('Triệu chứng đầy đủ không suy giảm', '0,6'),
        ('Triệu chứng dưới ngưỡng không suy giảm', '7,5'),
        ('TỔNG VẤN ĐỀ SKTT', '21,7'),
    ],
    header_bg='FFE6CC')
P('Hình 2. Tỷ lệ 12 tháng qua của vấn đề SKTT theo ngưỡng triệu chứng và suy giảm chức năng. Nguồn: Figure 2 báo cáo gốc, trang 21.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('DISC-5 đánh giá suy giảm chức năng do triệu chứng (theo yêu cầu DSM-5) trên bốn lĩnh vực: gia đình (vấn đề quan hệ với người chăm sóc, khó khăn khi dành thời gian với gia đình), bạn bè (khó khăn khi dành thời gian với peer), trường/làm việc, và căng thẳng cá nhân. Suy giảm ở nhiều hơn một lĩnh vực có thể được ghi nhận. Trong nhóm báo cáo suy giảm (n = 819), hai phần ba (67,0 %) báo cáo suy giảm ở lĩnh vực gia đình, trong khi gần một nửa báo cáo suy giảm liên quan bạn bè (47,0 %) và trường/làm việc (45,4 %) (Bảng 7).')

MakeTable(
    ['Lĩnh vực suy giảm', '%', 'n'],
    [
        ('Gia đình', '67,0', '549'),
        ('Bạn bè', '47,0', '384'),
        ('Trường hoặc việc làm', '45,4', '371'),
        ('Căng thẳng cá nhân', '34,6', '284'),
    ])
P('Bảng 7. Tỷ lệ 12 tháng qua của vấn đề SKTT theo lĩnh vực suy giảm. Gia trọng N = 819. Nguồn: Bảng 7 báo cáo gốc, trang 21.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H4('Rối loạn tâm thần (Mental disorders)')
P('Tiêu chí chẩn đoán DSM-5 cho bất kỳ rối loạn tâm thần nào được đáp ứng bởi 3,3 % VTN (n = 200), với dưới 1 % (0,6 %; n = 38) có 2 hoặc nhiều rối loạn trong 12 tháng qua. Không có khác biệt về tỷ lệ giữa nam (3,3 %) và nữ (3,4 %), hay giữa VTN trẻ hơn (2,9 %) và lớn hơn (3,9 %). Như Bảng 8 cho thấy, các rối loạn lo âu có tỷ lệ cao nhất (2,3 %).')

PageMarker('22')
MakeTable(
    ['Rối loạn tâm thần', '%', 'n'],
    [
        ('Các rối loạn lo âu *', '2,3', '135'),
        ('Rối loạn trầm cảm nặng', '0,9', '51'),
        ('ADHD', '0,5', '29'),
        ('PTSD', '0,3', '19'),
        ('Rối loạn hành vi', '0,2', '12'),
    ])
P('Bảng 8. Tỷ lệ 12 tháng qua của rối loạn tâm thần ở VTN 10–17 tuổi theo loại. * Tỷ lệ rối loạn lo âu cao hơn có ý nghĩa thống kê so với tất cả rối loạn khác. Nguồn: Bảng 8 báo cáo gốc, trang 22.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('Trong nhóm có bất kỳ rối loạn tâm thần nào, suy giảm ở lĩnh vực gia đình vẫn cao nhất với gần ba phần tư báo cáo suy giảm ở lĩnh vực này (Bảng 9).')

MakeTable(
    ['Lĩnh vực suy giảm', '%', 'n'],
    [
        ('Gia đình', '74,2', '149'),
        ('Bạn bè', '63,3', '127'),
        ('Trường hoặc việc làm', '64,1', '128'),
        ('Căng thẳng cá nhân', '51,8', '104'),
    ])
P('Bảng 9. Tỷ lệ 12 tháng qua của rối loạn tâm thần theo lĩnh vực suy giảm. Gia trọng N = 200. Nguồn: Bảng 9 báo cáo gốc, trang 22.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

H4('Hành vi tự sát và tự làm hại bản thân')
P('Ít VTN báo cáo hành vi tự sát trong 12 tháng qua. Trong toàn mẫu, 12 tháng qua: 1,4 % báo cáo ý nghĩ tự sát, 0,4 % báo cáo lập kế hoạch tự sát, 0,2 % báo cáo toan tự sát. Chỉ 1,6 % báo cáo đã từng toan tự sát trong đời. Tuy nhiên, như Bảng 10 cho thấy, trên 70 % các em báo cáo bất kỳ hành vi tự sát (ý nghĩ, kế hoạch, và/hoặc toan) trong 12 tháng qua đều đồng thời có vấn đề SKTT.')

MakeTable(
    ['Nhóm', '% Ý nghĩ tự sát 12m', 'n', '% Kế hoạch 12m', 'n', '% Toan tự sát 12m', 'n', '% Toan tự sát (từng)', 'n'],
    [
        ('Có vấn đề SKTT', '73,5', '61', '83,7', '19', '77,4', '8', '73,5', '70'),
        ('Có rối loạn tâm thần', '30,8', '26', '46,6', '11', '61,3', '7', '28,1', '27'),
    ])
P('Bảng 10. Hành vi tự sát ở VTN 10–17 tuổi. Gia trọng N: ý nghĩ tự sát 12 tháng = 83; kế hoạch 12 tháng = 23; toan tự sát 12 tháng = 11; toan tự sát từng = 96. Nguồn: Bảng 10 báo cáo gốc, trang 22.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('Tương tự, chỉ một tỷ lệ tương đối nhỏ (4,7 %) báo cáo đã từng tự làm hại bản thân, với chỉ hơn 1 % báo cáo tự làm hại trong 12 tháng qua (Bảng 11). Tuy nhiên, hơn ba phần tư (76,3 %) VTN tự làm hại trong 12 tháng qua có vấn đề SKTT.')

MakeTable(
    ['Nhóm', '% Tự làm hại 12m', 'n', '% Từng tự làm hại', 'n'],
    [
        ('Có vấn đề SKTT', '76,3', '50', '64,5', '182'),
        ('Có rối loạn tâm thần', '29,0', '19', '20,9', '59'),
    ])
P('Bảng 11. Tự làm hại bản thân ở VTN 10–17 tuổi. Gia trọng N: tự làm hại 12 tháng = 66; từng tự làm hại = 283. Nguồn: Bảng 11 báo cáo gốc, trang 22.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

PageMarker('23')
H3('Suy xét (Considerations)')

H4('Diễn giải (Interpretation)')
P('V-NAMHS tìm thấy SKTT kém là vấn đề sức khoẻ phổ biến ở VTN. Một trong năm VTN (21,7 %) có vấn đề SKTT trong 12 tháng qua, trong đó 1 trong 30 (3,3 %) đáp ứng tiêu chí cho rối loạn tâm thần. Về so sánh, có ít nghiên cứu đã sử dụng công cụ chẩn đoán để đo tỷ lệ rối loạn tâm thần trong quần thể VTN VN. Vượt qua các hạn chế phương pháp này là đặc điểm cốt lõi của V-NAMHS. Do khác biệt cơ bản về phương pháp, so sánh trực tiếp với nghiên cứu hiện có là khó khăn. Ví dụ, phát hiện 21,7 % VTN có vấn đề SKTT rơi vào khoảng 8–29 % của rà soát tài liệu Samuels et al. (2018). Tuy nhiên, so sánh có ý nghĩa giữa các công cụ khác nhau bị hạn chế, và không tìm được nghiên cứu nào có thể so sánh với phát hiện 3,3 % về tỷ lệ rối loạn tâm thần. Sự thiếu vắng nghiên cứu có thể so sánh khẳng định khoảng trống chứng cứ đáng kể mà V-NAMHS đã lấp đầy, cũng như tầm quan trọng của việc sử dụng công cụ toàn diện trên mẫu đại diện quốc gia.')

P('VTN báo cáo hành vi tự sát hoặc tự làm hại trong 12 tháng qua nhiều khả năng là những em có vấn đề SKTT. Tuy nhiên, tỷ lệ tổng thể hành vi tự sát và tự làm hại trong V-NAMHS hơi thấp hơn so với một số nghiên cứu trước. Ví dụ, Khảo sát Sức khoẻ Học sinh Toàn cầu (GSHS) 2019 tại VN tìm thấy 15,6 % VTN 13–17 tuổi báo cáo ý nghĩ tự sát trong 12 tháng qua, và 3,1 % báo cáo đã toan tự sát trong 12 tháng qua (WHO, MOH và MOET 2022). Tuy nhiên, GSHS dùng bộ câu hỏi tự quản trong bối cảnh trường học. Do đó có thể kỳ thị quanh hành vi tự sát góp phần giải thích tỷ lệ thấp hơn trong V-NAMHS, nơi câu hỏi về hành vi tự sát được phỏng vấn viên trực tiếp đặt. Hơn nữa, phát hiện GSHS cao hơn đáng kể so với các phát hiện liên quan khác tại VN. Ví dụ, Blum, Sudhinaraset và Emerson (2012) khảo sát 6.191 người 15–24 tuổi tại Hà Nội tìm thấy tỷ lệ 12 tháng của ý nghĩ tự sát là 2,3 %, với dưới 1 % báo cáo đã toan tự sát. Các tác giả nêu rằng câu hỏi nhạy cảm được tự quản bởi người tham gia; ngoài ra, câu hỏi ý nghĩ tự sát gộp chung ý nghĩ tự sát với tự làm hại, và kế hoạch tự sát không được bao gồm như danh mục riêng. Điều này có thể giải thích thêm tỷ lệ cao tương đối so với V-NAMHS.')

H4('Hạn chế (Limitations)')
P('Dù V-NAMHS được thiết kế chủ ý để giải quyết các hạn chế phương pháp của nghiên cứu hiện có, một số khía cạnh phương pháp của V-NAMHS có thể đã ảnh hưởng đến tỷ lệ báo cáo. Ví dụ, phỏng vấn được thực hiện trực tiếp bởi phỏng vấn viên được đào tạo. Có thể kỳ thị liên quan đến vấn đề SKTT và thiếu nhận thức chung về SKTT trong cộng đồng đã ảnh hưởng sự sẵn lòng tiết lộ thông tin của người tham gia, dù có đào tạo phỏng vấn viên toàn diện, yêu cầu riêng tư và tuyên bố dữ liệu ẩn danh. Vấn đề tương tự có thể có mặt với báo cáo hành vi tự sát và tự làm hại, vì các hành vi này thường bị kỳ thị, khiến VTN có thể miễn cưỡng báo cáo với phỏng vấn viên, dẫn đến tỷ lệ thấp hơn so với các nghiên cứu dùng bộ câu hỏi tự quản (Blum, Sudhinaraset and Emerson 2012, Hoàng Thế Hải và Bùi Thị Thanh Diệu 2021, UNICEF and ODI 2018, WHO, MOH và MOET 2022).')

PageMarker('24')
P('Một thách thức nữa liên quan đến định nghĩa rối loạn tâm thần theo DSM-5, là cơ sở của DISC-5. DSM-5 do Hiệp hội Tâm thần Hoa Kỳ phát triển và chủ yếu bao gồm tiêu chí chẩn đoán phương Tây (American Psychiatric Association 2013), có thể không tính đến khác biệt văn hoá về cách SKTT được mô tả, trải nghiệm hay thể hiện ở quần thể VN. Ví dụ, Kim và cộng sự (2019) tìm thấy triệu chứng cơ thể (somatic symptoms) dự báo triệu chứng lo âu và trầm cảm ở cả VTN VN lẫn VTN người Mỹ gốc Việt nhưng không ở VTN người Mỹ gốc Âu (Kim et al. 2019). Do đó có thể các triệu chứng trong DISC-5 (dựa trên DSM-5) không tính đủ triệu chứng cơ thể liên quan SKTT trong bối cảnh VN.')

P('Để giải quyết các thách thức này, đã có nhiều nỗ lực đáng kể để thích ứng DISC-5 với bối cảnh VN trong khuôn khổ yêu cầu chẩn đoán DSM-5. Bao gồm: dịch cẩn thận toàn bộ công cụ V-NAMHS (gồm DISC-5), dịch ngược, và rà soát dịch ngược bởi các nhóm NAMHS. Chỉnh sửa cũng được thực hiện dựa trên phản hồi của người tham gia các buổi đào tạo ban đầu và rà soát bản dịch của bác sĩ lâm sàng trong nước. Mục tiêu là đảm bảo ngôn ngữ của công cụ (ngữ pháp, thành ngữ, ví dụ về hành vi) được thích ứng đầy đủ với bối cảnh VN, đồng thời vẫn đo cùng khái niệm gốc. Thay đổi cũng được thực hiện sau nghiên cứu thí điểm (mô tả kỹ hơn trong Phụ lục 2) và tiếp tục qua 2019 và 2020. Cuối cùng, để hiểu tốt hơn sự tương tác tiềm năng giữa khác biệt văn hoá về SKTT và yêu cầu chẩn đoán DSM-5, cả "vấn đề SKTT" và "rối loạn tâm thần" đều được bao gồm trong báo cáo. Điều này cho phép hiểu toàn diện hơn về SKTT VTN VN, vượt xa yêu cầu chẩn đoán DSM-5.')

H4('Hàm ý (Implications)')
P('Tỷ lệ SKTT kém được V-NAMHS khám phá cho thấy SKTT là vấn đề y tế công cộng đòi hỏi sự chú ý của các nhà hoạch định chính sách và kế hoạch tại VN. Dữ liệu từ V-NAMHS cung cấp nền bằng chứng cho các chính sách và sáng kiến y tế như vậy. Ví dụ, đa số VTN hiện đang đi học, nghĩa là triển khai các chiến lược sàng lọc và quản lý cụ thể, lồng ghép với các hoạt động nâng cao SKTT trong trường học, có thể là một phương tiện giảm tỷ lệ và tác động của các triệu chứng này lên sức khoẻ và hạnh phúc VTN VN. Thêm nữa, lo âu có tỷ lệ cao nhất trong các vấn đề SKTT. Các chiến dịch nâng cao sức khoẻ có mục tiêu tập trung vào cải thiện hiểu biết về lo âu, nhận diện triệu chứng và cách tìm kiếm trợ giúp có thể tạo đường dẫn cải thiện SKTT cho một tỷ lệ lớn VTN VN có vấn đề SKTT.')

P('Ngoài ra, dù tỷ lệ báo cáo hành vi tự sát ở VTN thấp, vẫn cần sự chú ý chính sách và hành động để giảm thiểu nguy cơ khi câu hỏi là về sinh tử. Ý nghĩ tự sát thường xảy ra trong một khoảng thời gian trước khi kế hoạch được lập hoặc toan tự sát xảy ra (Wasserman et al. 2008). Do đó, phát hiện sớm và dịch vụ hỗ trợ cho những người có ý nghĩ tự sát — bao gồm thông qua nhận diện vấn đề SKTT — có thể giúp tránh toan tự sát sau này.')

P('Rộng hơn, phát hiện của V-NAMHS nhấn mạnh nhu cầu một chính sách SKTT quốc gia quan tâm đến VTN và có chiến lược cụ thể cho SKTT các em trong trường học và cộng đồng. Gần đây, Bộ Giáo dục và Đào tạo đã ban hành Quyết định 1442/QĐ-BGDĐT ngày 01/06/2022 (Ban hành Chương trình giáo dục sức khoẻ tâm thần cho trẻ em, học sinh giai đoạn 2022–2025) và Quyết định 2138/QĐ-BGDĐT ngày 03/08/2022 (Ban hành Kế hoạch giáo dục sức khoẻ tâm thần cho trẻ em, học sinh giai đoạn 2022–2025) (Bộ Giáo dục và Đào tạo 2022a, 2022b). Dù các quyết định này thể hiện sự công nhận rộng rãi về SKTT VTN trong ngành giáo dục, vẫn cần một khung chính sách toàn diện cấp quốc gia để giải quyết nhu cầu SKTT của VTN, bao gồm cả các em trong và ngoài trường học, và bao gồm sự tham gia của ngành y tế cùng các ngành liên quan khác. Hành động tích cực cho SKTT VTN có tiềm năng mang lại lợi ích hiện tại (giảm triệu chứng, cải thiện khả năng hoạt động tối ưu hàng ngày), lợi ích dài hạn (giảm hệ quả xấu ở tuổi trưởng thành), cũng như lợi ích cho thế hệ kế tiếp bằng việc có phụ huynh SKTT tốt và lớn lên trong xã hội có chính sách SKTT được thiết lập (Patton et al. 2016).')

doc.save(OUT)
print(f'Part 3 (Mental Health chapter) saved.')
