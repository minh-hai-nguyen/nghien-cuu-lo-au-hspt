# -*- coding: utf-8 -*-
"""Match author citations in 10/04/2026 reports → canonical papers in 02_Papers-goc/."""
import os, sys, io, re, json, glob
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Main 10/04 report (the primary intervention report)
MAIN_REPORT = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 10042026 v2.docx')

# Load canonical index
with open(os.path.join(ROOT, '02_Papers-goc', 'canonical_index.json'), encoding='utf-8') as f:
    canon = json.load(f)

# Build descriptor → canonical_id lookup
descr_to_id = {}
for cid, meta in canon.items():
    descr = meta.get('descriptor', '').lower()
    descr_to_id[descr] = cid

# Read report
d = Document(MAIN_REPORT)
report_text = '\n'.join(p.text for p in d.paragraphs if p.text.strip())
for t in d.tables:
    for r in t.rows:
        for c in r.cells:
            report_text += ' ' + c.text

print(f'Report: {os.path.basename(MAIN_REPORT)}')
print(f'Text length: {len(report_text):,} chars')
print('='*70)

# Extract REAL author + year citations
# Focus on meaningful ones — look for patterns like "(Author, Year)" or "Author et al., Year" or "Author (Year)"
citation_patterns = [
    r'\(([A-Z][a-zà-ỹ\-]+(?:\s+et\s+al\.?)?(?:\s+và\s+cộng\s+sự)?),?\s+(\d{4})\)',  # (Author, 2023)
    r'([A-Z][a-zà-ỹ\-]+(?:\s+et\s+al\.?)?)\s*\((\d{4})\)',  # Author (2023)
    r'([A-Z][a-zà-ỹ\-]+)\s+và\s+cộng\s+sự\s+(\d{4})',  # Author và cộng sự 2023
    r'([A-Z][a-zà-ỹ\-]+)\s+et\s+al\.?\s+(\d{4})',  # Author et al. 2023
]

raw_cites = set()
for pat in citation_patterns:
    for m in re.finditer(pat, report_text):
        author = m.group(1).strip().rstrip('.')
        year = m.group(2)
        if 2000 <= int(year) <= 2026:
            raw_cites.add((author, year))

# Filter common noise (not author names)
NOISE = {'bảng', 'hình', 'chương', 'mục', 'phần', 'trang', 'chỉ', 'báo', 'đề',
         'hoa kỳ', 'hàn quốc', 'việt nam', 'đông nam á', 'ireland tăng', 'ireland',
         'korea', 'europe', 'asean', 'nma', 'cbt', 'ahmad', 'namhs', 'gbd',
         'jama', 'bmj', 'nhs', 'ajp', 'interv', 'epipsychsci', 'jaacap',
         'mục tiêu', 'countries', 'long an', 'an giang', 'hải phòng', 'vĩnh long',
         'tph', 'hmu', 'y học', 'y khoa', 'y học cộng đồng', 'thuốc', 'medicine',
         'lancet', 'kids', 'worldwide', 'march', 'mental', 'schools', 'mindfulness'}
citations = sorted([(a, y) for a, y in raw_cites if a.lower() not in NOISE])

print(f'\nClean citations extracted: {len(citations)}')
for a, y in citations:
    print(f'  {a} {y}')

# Now match each against canonical papers
print(f'\n{"="*70}')
print('MATCHING CITATIONS → CANONICAL PAPERS')
print('='*70)

# Build author lookup: canonical_id → set of (author_surname_lowercase, year)
id_to_authors = {}
for cid, meta in canon.items():
    descr = meta.get('descriptor', '')
    # descriptor format: "Author_et_al_YYYY_Country" or "Hoa_2024_Frontiers_HaNoi"
    m = re.match(r'^([A-Za-z\-]+)', descr)
    surname = m.group(1).lower() if m else ''
    year_m = re.search(r'(\d{4})', descr)
    year = year_m.group(1) if year_m else ''
    id_to_authors[cid] = (surname, year, descr)

matched = []
unmatched = []
for author, year in citations:
    au_low = author.lower().replace(' et al.', '').replace(' et al', '').strip()
    # Match attempt
    found = None
    for cid, (surname, cyear, descr) in id_to_authors.items():
        if (au_low == surname or au_low in surname or surname in au_low) and year == cyear:
            found = (cid, descr)
            break
    if found:
        matched.append((author, year, found[0], found[1]))
    else:
        unmatched.append((author, year))

print(f'\nMATCHED ({len(matched)}):')
for au, yr, cid, descr in matched:
    pdf = canon[cid].get('pdf', '?')
    print(f'  {au} {yr}  →  {cid}  [{descr}]  (PDF: {pdf})')

print(f'\nUNMATCHED ({len(unmatched)}) — need to check:')
for au, yr in unmatched:
    print(f'  {au} {yr}')

# Save report
out_data = {
    'report': os.path.basename(MAIN_REPORT),
    'total_citations': len(citations),
    'matched': [{'author': a, 'year': y, 'canonical_id': c, 'descriptor': d} for a, y, c, d in matched],
    'unmatched': [{'author': a, 'year': y} for a, y in unmatched],
}
OUT = os.path.join(os.path.dirname(__file__), 'citations_10042026_match.json')
with open(OUT, 'w', encoding='utf-8') as f:
    json.dump(out_data, f, ensure_ascii=False, indent=2)
print(f'\nReport saved: {OUT}')
