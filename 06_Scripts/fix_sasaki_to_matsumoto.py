# -*- coding: utf-8 -*-
"""
Fix: QT045 tác giả chính là Matsumoto et al. 2024, KHÔNG PHẢI Sasaki.
Actual first author: Kazuki Matsumoto, PhD (Kagoshima University Hospital).
10 co-authors: Hamatani, Shiga, Iiboshi, Kasai, Kimura, Yokota, Watanabe, Kubo, Nakamura.
No author named Sasaki in this paper.

Fix scope (per Nguyên tắc 7 — preserve old snapshots):
- v5.5 (latest report)
- Tom-tat-tung-bai/QT045
- canonical_index.json
- Bai_goc folder (README + PDF rename + QT050 cross-ref)
- Rename PDF files
"""
import os, sys, io, re, json, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# ============================================================
# Text replacements for docx content
# ============================================================
REPLACEMENTS = [
    ('Sasaki et al. 2024', 'Matsumoto et al. 2024'),
    ('Sasaki et al 2024', 'Matsumoto et al 2024'),
    ('Sasaki 2024', 'Matsumoto 2024'),
    ('Sasaki, N., et al. (2024)', 'Matsumoto, K., Hamatani, S., Shiga, K., et al. (2024)'),
    ('Sasaki et al.', 'Matsumoto et al.'),
    # In tables / references
    ('Bài 51 — Sasaki, N., et al. (2024)', 'Bài 51 — Matsumoto, K., Hamatani, S., Shiga, K., et al. (2024)'),
    ('Sasaki_Japan_iCBT_JMIR_2024', 'Matsumoto_Japan_iCBT_JMIR_2024'),
    ('QT045_Sasaki', 'QT045_Matsumoto'),
    # Keep Sasaki intact if it's in a different context (Wen_et_al_2020 reference — check separately)
]

def fix_docx(path, replacements):
    """Fix all Sasaki → Matsumoto in a docx file."""
    d = Document(path)
    changes = 0
    for p in d.paragraphs:
        t = p.text
        new_t = t
        for find, rep in replacements:
            if find in new_t:
                new_t = new_t.replace(find, rep)
        if new_t != t:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_t
            changes += 1
    for tbl in d.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    t = p.text
                    new_t = t
                    for find, rep in replacements:
                        if find in new_t:
                            new_t = new_t.replace(find, rep)
                    if new_t != t:
                        for run in p.runs:
                            run.text = ''
                        if p.runs:
                            p.runs[0].text = new_t
                        changes += 1
    if changes > 0:
        d.save(path)
    return changes

# ============================================================
# 1. Fix latest report v5.5
# ============================================================
v55 = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-5.docx')
if os.path.exists(v55):
    n = fix_docx(v55, REPLACEMENTS)
    print(f'[1] v5.5 report: {n} changes')

# ============================================================
# 2. Fix Tom-tat-tung-bai/QT045 (content + rename file)
# ============================================================
old_tt = os.path.join(ROOT, 'Tom-tat-tung-bai', 'QT045_Sasaki_Japan_iCBT_JMIR_2024.docx')
new_tt = os.path.join(ROOT, 'Tom-tat-tung-bai', 'QT045_Matsumoto_Japan_iCBT_JMIR_2024.docx')
if os.path.exists(old_tt):
    n = fix_docx(old_tt, REPLACEMENTS)
    os.rename(old_tt, new_tt)
    print(f'[2] Tom-tat QT045: {n} changes + renamed → QT045_Matsumoto_Japan_iCBT_JMIR_2024.docx')

# ============================================================
# 3. Fix canonical_index.json
# ============================================================
idx_path = os.path.join(ROOT, '02_Papers-goc', 'canonical_index.json')
with open(idx_path, encoding='utf-8') as f:
    canon = json.load(f)

if 'QT045' in canon:
    old = canon['QT045']
    canon['QT045'] = {
        'id': 'QT045',
        'descriptor': 'Matsumoto_Japan_iCBT_JMIR_2024',
        'summary': 'QT045_Matsumoto_Japan_iCBT_JMIR_2024.docx',
        'translation': old.get('translation', '').replace('Sasaki', 'Matsumoto'),
        'pdf': 'QT045_Matsumoto_Japan_iCBT_JMIR_2024.pdf',
        'pdf_folder': old.get('pdf_folder', 'The-gioi_Khac'),
    }
    with open(idx_path, 'w', encoding='utf-8') as f:
        json.dump(canon, f, ensure_ascii=False, indent=2)
    print(f'[3] canonical_index.json: QT045 updated → Matsumoto')

# ============================================================
# 4. Rename PDF in 02_Papers-goc/
# ============================================================
old_pdf = os.path.join(ROOT, '02_Papers-goc', 'The-gioi_Khac', 'QT045_Sasaki_Japan_iCBT_JMIR_2024.pdf')
new_pdf = os.path.join(ROOT, '02_Papers-goc', 'The-gioi_Khac', 'QT045_Matsumoto_Japan_iCBT_JMIR_2024.pdf')
if os.path.exists(old_pdf):
    os.rename(old_pdf, new_pdf)
    print(f'[4] PDF renamed: QT045_Sasaki → QT045_Matsumoto')

# ============================================================
# 5. Fix + rename PDF in Bai_goc folder
# ============================================================
old_bai = os.path.join(ROOT, 'Bai_goc_BaoCao_CanThiep_10042026', 'Bai_51__QT045_Sasaki_Japan_iCBT_JMIR_2024.pdf')
new_bai = os.path.join(ROOT, 'Bai_goc_BaoCao_CanThiep_10042026', 'Bai_51__QT045_Matsumoto_Japan_iCBT_JMIR_2024.pdf')
if os.path.exists(old_bai):
    os.rename(old_bai, new_bai)
    print(f'[5] Bai_goc PDF renamed: Sasaki → Matsumoto')

# ============================================================
# 6. Fix README.md in Bai_goc folder
# ============================================================
readme_path = os.path.join(ROOT, 'Bai_goc_BaoCao_CanThiep_10042026', 'README.md')
if os.path.exists(readme_path):
    with open(readme_path, encoding='utf-8') as f:
        content = f.read()
    old_count = content.count('Sasaki')
    for find, rep in REPLACEMENTS:
        content = content.replace(find, rep)
    # Also fix raw filename references
    content = content.replace('Bai_51__QT045_Sasaki', 'Bai_51__QT045_Matsumoto')
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(content)
    print(f'[6] README.md: {old_count} Sasaki → Matsumoto')

# ============================================================
# 7. Fix QT050 enhanced file (cross-references Sasaki)
# ============================================================
qt050_path = os.path.join(ROOT, 'Bai_goc_BaoCao_CanThiep_10042026',
                          'Bai_57_QT050_Qiaochu_Wang_Mobile_CBT_2025_BAN_DICH_CHI_TIET.docx')
if os.path.exists(qt050_path):
    n = fix_docx(qt050_path, REPLACEMENTS)
    print(f'[7] QT050 enhanced: {n} changes (cross-ref Sasaki → Matsumoto)')

# ============================================================
# 8. Fix QT050 original Tom-tat if exists
# ============================================================
qt050_tt = os.path.join(ROOT, 'Tom-tat-tung-bai', 'QT050_Qiaochu_MobileCBT_CPP_2025_ABSTRACT.docx')
if os.path.exists(qt050_tt):
    n = fix_docx(qt050_tt, REPLACEMENTS)
    print(f'[8] QT050 Tom-tat: {n} changes')

# ============================================================
# 9. Verify no Sasaki remains in key files
# ============================================================
print('\n' + '='*70)
print('VERIFICATION — "Sasaki" should be 0 in fixed files:')
check_files = [
    v55,
    new_tt,
    readme_path,
    qt050_path,
    qt050_tt,
]
for path in check_files:
    if not os.path.exists(path):
        print(f'  [SKIP] {os.path.basename(path)} not found')
        continue
    if path.endswith('.docx'):
        d = Document(path)
        txt = '\n'.join(p.text for p in d.paragraphs)
        for t in d.tables:
            for r in t.rows:
                for c in r.cells:
                    txt += ' ' + c.text
        count = txt.count('Sasaki')
    elif path.endswith('.md'):
        with open(path, encoding='utf-8') as f:
            txt = f.read()
        count = txt.count('Sasaki')
    else:
        count = -1
    status = 'OK' if count == 0 else f'STILL {count} hits'
    print(f'  {os.path.basename(path)}: {status}')

# Check Matsumoto is present
print('\nVERIFICATION — "Matsumoto" should be present:')
for path in check_files:
    if not os.path.exists(path): continue
    if path.endswith('.docx'):
        d = Document(path)
        txt = '\n'.join(p.text for p in d.paragraphs)
        count = txt.count('Matsumoto')
    elif path.endswith('.md'):
        with open(path, encoding='utf-8') as f:
            count = f.read().count('Matsumoto')
    else:
        count = -1
    print(f'  {os.path.basename(path)}: {count} hits')
