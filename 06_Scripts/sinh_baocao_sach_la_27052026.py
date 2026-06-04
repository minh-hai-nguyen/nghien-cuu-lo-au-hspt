# -*- coding: utf-8 -*-
"""Sinh bao cao sach - chi LA, dang bang ngan gon de sua nhanh.
Bo metadata + watermark.
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA_REF = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_26052026.docx')  # original to compute pages
OUT = os.path.join(ROOT, 'Luận án TS', '08_BaoCao_Sach_27052026.docx')

RED = RGBColor(192, 0, 0)
GREEN = RGBColor(0, 112, 0)
GRAY = RGBColor(80, 80, 80)


# Estimate page from paragraph index using simple heuristic
# Based on LA layout: ~12 paragraphs/page average
def est_page(para_idx):
    if para_idx < 100:  # front matter
        return max(1, para_idx // 8)
    if para_idx < 300:  # Chương 1 Tổng quan
        return 10 + (para_idx - 100) // 10
    if para_idx < 600:  # Chương 1 + Chương 2
        return 30 + (para_idx - 300) // 10
    if para_idx < 900:  # Chương 2 + Chương 3
        return 60 + (para_idx - 600) // 10
    if para_idx < 1200:  # Chương 3 Kết quả
        return 90 + (para_idx - 900) // 10
    if para_idx < 1400:  # Bàn luận + TLTK begin
        return 120 + (para_idx - 1200) // 8
    return 145 + (para_idx - 1400) // 8


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(2); s.paragraph_format.line_spacing = 1.15
    for sec in d.sections:
        sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(1.8); sec.right_margin = Cm(1.5)
    # Bo metadata
    cp = d.core_properties
    cp.author = ''
    cp.title = ''
    cp.subject = ''
    cp.keywords = ''
    cp.comments = ''
    cp.last_modified_by = ''
    return d


def H(d, text, level=1):
    h = d.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
    return h


def P(d, text, bold=False, italic=False, size=11, color=None):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color is not None:
        r.font.color.rgb = color
    return p


def add_fix_table(d, fixes):
    """Add a table with cols: # | Trang~ | Đoạn | CŨ | MỚI."""
    t = d.add_table(rows=1, cols=5)
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    headers = ['#', 'Trang~', 'Đoạn', 'CŨ (sai)', 'MỚI (đã sửa)']
    widths = [Cm(0.7), Cm(1.2), Cm(1.3), Cm(7.5), Cm(7.5)]
    for i, (h, w) in enumerate(zip(headers, widths)):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
        hdr[i].width = w
    for i, (para_idx, old, new) in enumerate(fixes, 1):
        row = t.add_row().cells
        row[0].text = str(i)
        row[1].text = f"~{est_page(para_idx)}"
        row[2].text = str(para_idx)
        # Cũ
        row[3].text = ''
        p = row[3].paragraphs[0]
        r = p.add_run(old)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RED
        # Mới
        row[4].text = ''
        p = row[4].paragraphs[0]
        r = p.add_run(new)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = GREEN
        # Set widths on data cells
        for j, w in enumerate(widths):
            row[j].width = w


# ============================================================
# DATA - 15 thay doi can sua trong LA (khong tinh global term)
# ============================================================
FIXES_DETAILED = [
    # (para_idx, OLD text, NEW text)
    (267,
     'tại 2 trường THPT ở huyện Y Đ - vùng bán đô thị, gần thành phố Thanh Hóa',
     'tại 2 trường THPT ở huyện Y Đ - huyện bán nông thôn, gần thành phố Thanh Hóa'),
    (268,
     'có 49% học sinh lo âu, trong đó có 7,7% loại nhẹ, 24,5% loại vừa, 8,1% loại nặng, 4,6% loại rất nặng',
     'có 49,0% học sinh lo âu, trong đó có 11,2% loại nhẹ, 25,1% loại vừa, 8,1% loại nặng, 4,6% loại rất nặng'),
    (269,
     'phòng ngừa rối loạn lo âu cho học sinh ở những địa bàn dân cư bán đô thị như Yên Định, Thanh Hóa',
     'phòng ngừa rối loạn lo âu cho học sinh ở những địa bàn dân cư bán nông thôn như Yên Định, Thanh Hóa'),
    (271,
     'Công trình của Lê Minh T., Nguyễn Đăng K., Ngô Anh V. (2025)',
     'Công trình của Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh (2025)'),
    (315,
     '61,2% theo Lê Minh T., 2025',
     '61,2% theo Nguyễn Đăng Khoa và cs., 2025'),
    (1368,
     'Lê Minh Thi, Nguyễn Đăng Khoa, & Ngô Anh Vinh (2025).',
     'Nguyễn Đăng Khoa, Lê Minh Thi, & Ngô Anh Vinh (2025).'),
    (174,
     'M. Sakia và cs.(2023)',
     'A.M. Saikia và cs. (2023)'),
    (169,
     'Q. Xu và cs (2022) đưa ra tỷ lệ chênh lệch nam lo âu hơn nữ là 10,11% nam/9,66% nữ',
     'Q. Xu và cs. (2021) đưa ra tỷ lệ chênh lệch nam lo âu hơn nữ là 10,11% nam/9,66% nữ'),
    (235,
     'Xu và cs (2022) khảo sát ở Trung Quốc trong đại dịch Covid-19',
     'Xu và cs. (2021) khảo sát ở Trung Quốc trong đại dịch Covid-19'),
    (238,
     '… Chongjian Wang, Cuiping Wu (2022) về Tỷ lệ và các yếu tố nguy cơ…',
     '… Chongjian Wang, Cuiping Wu (2021) về Tỷ lệ và các yếu tố nguy cơ…'),
    (320,
     '… (A.M. Saikia và cs., 2023), (Q. Xu và cs., 2022)',
     '… (A.M. Saikia và cs., 2023), (Q. Xu và cs., 2021)'),
    (169,
     '(LA tổng quát, LA xã hội và LA chia ly)',
     '(LA lan tỏa, LA xã hội và LA chia ly)'),
    (1113,
     'các mô hình đơn (RLLA tổng quát, chia ly và xã hội) đều có chỉ số phù hợp tốt',
     'các mô hình đơn (RLLA lan tỏa, chia ly và xã hội) đều có chỉ số phù hợp tốt'),
    (1115,
     'β → LATQ = 0,215; β → LAXH = 0,253',
     'β → LALT = 0,215; β → LAXH = 0,253'),
    (1324,
     '... LATQ ...',
     '... LALT ...'),
]


# ============================================================
# BUILD DOC
# ============================================================
d = doc_init()

H(d, 'BÁO CÁO SẠCH CÁC ĐIỂM CẦN SỬA TRONG LUẬN ÁN', 1)
P(d, '27/05/2026 — File hiện tại: 01_LuanAn_v3_1_FixCoping_26052026.docx', italic=True, size=10)
P(d, 'File đã sửa sẵn (chữ đỏ): 01_LuanAn_v3_4_FixAbbrev_27052026.docx', italic=True, size=10)
d.add_paragraph()

P(d, 'Báo cáo gồm 2 phần:', bold=True)
P(d, '   • Phần A — 15 điểm sửa cụ thể (tọa độ trang ~ + đoạn + nội dung cũ/mới).')
P(d, '   • Phần B — Sửa hàng loạt theo Find/Replace (thuật ngữ "tổng quát" → "lan tỏa").')
d.add_paragraph()

# =========================
# PHAN A
# =========================
H(d, 'A. CÁC ĐIỂM SỬA CỤ THỂ', 1)
P(d, '(Trang ước lượng giả định ~12 đoạn/trang. Mở file LA, dùng tổ hợp phím Ctrl+G để nhảy tới số đoạn cụ thể, hoặc Ctrl+F dán chuỗi "CŨ" để định vị nhanh.)', italic=True, size=10)
d.add_paragraph()

H(d, 'A.1. VN017 Nguyễn Danh Lâm 2022 — Sửa địa lý + số liệu', 2)
add_fix_table(d, FIXES_DETAILED[0:3])
d.add_paragraph()
P(d, 'Lý do: Yên Định là huyện thuần nông thôn (KHÔNG phải bán đô thị). Tỷ lệ lo âu nhẹ và vừa bị lẫn cột với hàng Stress khi sao chép số liệu — đã đối chiếu Bảng PDF gốc DOI:10.51298/vmj.v516i1.2948.', italic=True, size=10)
d.add_paragraph()

H(d, 'A.2. VN018 An Giang 2025 — Sửa thứ tự tác giả', 2)
add_fix_table(d, FIXES_DETAILED[3:6])
d.add_paragraph()
P(d, 'Lý do: Trong PDF gốc DOI:10.51298/vmj.v549i1.13506, tác giả thứ nhất là "Nguyễn Đăng Khoa", thứ hai là "Lê Minh Thi", thứ ba là "Ngô Anh Vinh". 3 vị trí cần sửa nhất quán: in-text giới thiệu công trình, citation in-line trong danh sách tỷ lệ, và TLTK cuối luận án.', italic=True, size=10)
d.add_paragraph()

H(d, 'A.3. Sửa chính tả tên tác giả', 2)
add_fix_table(d, FIXES_DETAILED[6:7])
d.add_paragraph()
P(d, 'Lý do: "Sakia" thiếu chữ "i" so với tên đầy đủ "Anku M. Saikia" trên PDF gốc DOI:10.4103/ijcm.ijcm_614_21.', italic=True, size=10)
d.add_paragraph()

H(d, 'A.4. Sửa năm xuất bản Xu et al. (2022 → 2021)', 2)
add_fix_table(d, FIXES_DETAILED[7:11])
d.add_paragraph()
P(d, 'Lý do: TLTK đoạn 1545 ghi đúng năm 2021 (Journal of Affective Disorders, vol. 288, 2021, DOI:10.1016/j.jad.2021.03.080). 4 vị trí in-text trong body trước đây ghi sai năm 2022. Sau khi sửa, in-text và TLTK đã khớp nhau.', italic=True, size=10)
d.add_paragraph()

H(d, 'A.5. Sửa viết tắt còn sót sau khi đổi "tổng quát" → "lan tỏa"', 2)
add_fix_table(d, FIXES_DETAILED[11:15])
d.add_paragraph()
P(d, 'Lý do: Sau khi đổi global "rối loạn lo âu tổng quát" → "rối loạn lo âu lan tỏa", 4 vị trí còn sót viết tắt "LA tổng quát" / "RLLA tổng quát" / "LATQ" (không bao gồm tiền tố RL).', italic=True, size=10)
d.add_paragraph()

# =========================
# PHAN B
# =========================
H(d, 'B. SỬA HÀNG LOẠT THEO FIND/REPLACE', 1)
P(d, 'Mở Word > Ctrl+H để dùng tính năng Find and Replace. Sửa các chuỗi sau (đảm bảo check vào "Match case" khi cần):', italic=True, size=10)
d.add_paragraph()

# Find-replace table
t = d.add_table(rows=1, cols=4)
t.style = 'Light Grid Accent 1'
headers_b = ['#', 'FIND (cũ)', 'REPLACE (mới)', 'Số chỗ ước tính']
widths_b = [Cm(0.7), Cm(7.0), Cm(7.0), Cm(2.5)]
for i, (h, w) in enumerate(zip(headers_b, widths_b)):
    t.rows[0].cells[i].text = ''
    p = t.rows[0].cells[i].paragraphs[0]
    r = p.add_run(h); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    t.rows[0].cells[i].width = w

bulk_fixes = [
    ('rối loạn lo âu tổng quát', 'rối loạn lo âu lan tỏa', '53'),
    ('Rối loạn lo âu tổng quát', 'Rối loạn lo âu lan tỏa', '~5'),
    ('Rối loạn Lo âu Tổng quát', 'Rối loạn Lo âu Lan tỏa', '~3'),
    ('lo âu tổng quát', 'lo âu lan tỏa', '82'),
    ('Lo âu tổng quát', 'Lo âu lan tỏa', '~3'),
    ('Lo âu Tổng quát', 'Lo âu Lan tỏa', '~2'),
    ('RLLATQ', 'RLLALT', '70'),
    ('RLLA tổng quát', 'RLLA lan tỏa', '1'),
    ('LA tổng quát', 'LA lan tỏa', '1'),
    ('LATQ', 'LALT', '2'),
]
for i, (old, new, n) in enumerate(bulk_fixes, 1):
    row = t.add_row().cells
    row[0].text = str(i)
    row[1].text = ''
    p = row[1].paragraphs[0]
    r = p.add_run(old); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True; r.font.color.rgb = RED
    row[2].text = ''
    p = row[2].paragraphs[0]
    r = p.add_run(new); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True; r.font.color.rgb = GREEN
    row[3].text = ''
    p = row[3].paragraphs[0]
    r = p.add_run(n); r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    for j, w in enumerate(widths_b):
        row[j].width = w

d.add_paragraph()
P(d, 'Lý do: DSM-5 và ICD-11 bản dịch tiếng Việt chuẩn dùng "Rối loạn lo âu lan tỏa" cho Generalized Anxiety Disorder (KHÔNG dùng "tổng quát"). RLLATQ là mã biến SEM tương ứng, đổi sang RLLALT để khớp tên thuật ngữ.', italic=True, size=10)
P(d, 'Tổng cộng ước tính ~222 chỗ. Khuyến nghị làm thứ tự từ chuỗi dài đến chuỗi ngắn để tránh sửa chéo (đã sắp xếp đúng thứ tự trong bảng trên).', italic=True, size=10)
d.add_paragraph()

# =========================
# TONG KET
# =========================
H(d, 'C. TÓM TẮT', 1)
P(d, 'Tổng số điểm sửa: ~237 (15 sửa cụ thể + ~222 chỗ qua Find/Replace).', bold=True)
P(d, 'File LA đã sửa sẵn (chữ đỏ + đậm để dễ nhìn): 01_LuanAn_v3_4_FixAbbrev_27052026.docx — chỉ cần mở lên là thấy toàn bộ thay đổi.')
P(d, 'Nếu muốn sửa thủ công từ bản gốc, theo Phần A (15 chỗ cụ thể) + Phần B (Find/Replace hàng loạt).')
d.add_paragraph()


# ============================================================
# REMOVE WATERMARK + HEADER/FOOTER (any)
# ============================================================
for sec in d.sections:
    for hdr in [sec.header, sec.first_page_header, sec.even_page_header]:
        # Remove pict elements (watermark)
        for elem in hdr._element.iter():
            if elem.tag == qn('w:pict') or elem.tag == qn('w:object'):
                elem.getparent().remove(elem)
        # Clear any text in headers
        for p in hdr.paragraphs:
            for r in p.runs:
                r.text = ''
    for ftr in [sec.footer, sec.first_page_footer, sec.even_page_footer]:
        for elem in ftr._element.iter():
            if elem.tag == qn('w:pict') or elem.tag == qn('w:object'):
                elem.getparent().remove(elem)
        for p in ftr.paragraphs:
            for r in p.runs:
                r.text = ''


d.save(OUT)
print(f"Da luu: {OUT}")
print(f"Paragraphs: {len(d.paragraphs)}, Tables: {len(d.tables)}, Size: {os.path.getsize(OUT)//1024}KB")
