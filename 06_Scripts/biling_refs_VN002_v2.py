# -*- coding: utf-8 -*-
"""Re-add VN translations for ALL English references (the first script failed)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor

OUT = '03_Ban-dich/VN002_VNAMHS_2022_National_FULL.docx'

# Prefix-matched entries (English refs only — VN refs already done)
# Use first 50 chars of EN ref as prefix matcher
BILING = [
    ('American Psychiatric Association. 2013. "Diagnostic',
     'Hiệp hội Tâm thần Hoa Kỳ. 2013. "Sổ tay Chẩn đoán và Thống kê Rối loạn Tâm thần: Tái bản thứ Năm, DSM-5." Washington DC: Nhà xuất bản Hiệp hội Tâm thần Hoa Kỳ.'),
    ('Bitsko, R., H.R. Adams, J. Holbrook',
     'Bitsko R. và cộng sự. 2019. "2.50 Bảng Phỏng vấn Chẩn đoán cho Trẻ em, Phiên bản 5 (DISC-5): Phát triển và Thẩm định các Module ADHD và Rối loạn Tic." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 58(10): Phụ lục.'),
    ('Blum, R., M. Sudhinaraset and M.R. Emerson',
     'Blum R., Sudhinaraset M. và Emerson M.R. 2012. "Vị thành niên có nguy cơ: Ý nghĩ và Hành vi Tự sát tại Việt Nam, Trung Quốc và Đài Loan." Tạp chí Sức khoẻ Vị thành niên 50(3):S37–S44.'),
    ('Canino, G. and M. Alegria',
     'Canino G. và Alegria M. 2008. "Chẩn đoán tâm thần — Có tính phổ quát hay tuỳ thuộc văn hoá?" Tạp chí Tâm lý học và Tâm thần học Trẻ em 49(3):237–50.'),
    ('Colizzi, M., A. Lasalvia',
     'Colizzi M., Lasalvia A. và Ruggeri M. 2020. "Phòng ngừa và Can thiệp Sớm trong Sức khoẻ Tâm thần Thanh thiếu niên: Đã đến lúc cho Mô hình Chăm sóc Đa ngành và Xuyên Chẩn đoán?" Tạp chí Quốc tế về Hệ thống Sức khoẻ Tâm thần 14.'),
    ('COVID-19 Mental Disorders Collaborators',
     'Nhóm Cộng tác Rối loạn Tâm thần COVID-19. 2021. "Tỷ lệ và Gánh nặng Toàn cầu của Rối loạn Trầm cảm và Lo âu tại 204 Quốc gia và Vùng lãnh thổ năm 2020 do Đại dịch COVID-19." The Lancet 398(10312):1700–12.'),
    ('Demyttenaere, K. et al. 2004',
     'Demyttenaere K. và cộng sự. 2004. "Tỷ lệ, Mức độ nghiêm trọng và Nhu cầu Điều trị chưa được Đáp ứng đối với Rối loạn Tâm thần trong các Khảo sát Sức khoẻ Tâm thần Thế giới của WHO." JAMA 291(21):2581–90.'),
    ('Erskine, H.E. et al. 2015',
     'Erskine H.E. và cộng sự. 2015. "Một Gánh nặng cho Tâm trí Trẻ: Gánh nặng Toàn cầu của Rối loạn Tâm thần và Sử dụng Chất ở Trẻ em và Thanh niên." Y học Tâm lý 45(7):1551–63.'),
    ('Erskine, H.E. et al. 2016',
     'Erskine H.E. và cộng sự. 2016. "Kết quả dài hạn của Rối loạn Tăng động Giảm Chú ý và Rối loạn Hành vi: Tổng quan Hệ thống và Phân tích Meta." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 55(10):841–50.'),
    ('Erskine, H.E., S. Blondell, M. Enright',
     'Erskine H.E. và cộng sự (bao gồm Vũ Mạnh Lợi, Đào Thị Khánh Hoa, Nguyễn Đức Vinh từ Việt Nam). 2021. "Đo lường Tỷ lệ Rối loạn Tâm thần ở Vị thành niên tại Kenya, Indonesia và Việt Nam: Đề cương Nghiên cứu cho các Khảo sát Sức khoẻ Tâm thần Vị thành niên Quốc gia." Tạp chí Sức khoẻ Vị thành niên.'),
    ('Ferrari, A.J. et al. 2013',
     'Ferrari A.J. và cộng sự. 2013. "Mô hình hoá Dịch tễ học Rối loạn Trầm cảm Nặng: Ứng dụng cho Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2010." PLOS One 8(7): e69637.'),
    ('GBD 2019 Mental Disorders Collaborators',
     'Nhóm Cộng tác GBD 2019 Rối loạn Tâm thần. 2022. "Gánh nặng Toàn cầu, Khu vực và Quốc gia của 12 Rối loạn Tâm thần tại 204 Quốc gia và Vùng lãnh thổ, 1990–2019: Phân tích Hệ thống cho Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019." The Lancet Psychiatry 9(2):137–50.'),
    ('General Statistical Office. 2020',
     '(Song ngữ gốc — không cần dịch bổ sung. Tác giả: Tổng cục Thống kê Việt Nam.)'),
    ('Hafekost, J. et al. 2016',
     'Hafekost J. và cộng sự. 2016. "Phương pháp của Young Minds Matter: Khảo sát Sức khoẻ Tâm thần và Hạnh phúc Trẻ em và Vị thành niên Úc lần thứ Hai." Tạp chí Tâm thần học Úc và New Zealand 50(9):866–75.'),
    ('Hafstad, G.S. and E. Augusti',
     'Hafstad G.S. và Augusti E. 2021. "Một Thế hệ Mất mát? COVID-19 và Sức khoẻ Tâm thần Vị thành niên." The Lancet Psychiatry 8(8):640–41.'),
    ('Hoang M.D., T.T. Lam, A. Dao and B. Weiss',
     'Hoàng Minh Đặng, Lâm T.T., Đào A. và Weiss B. 2020. "Kiến thức về Sức khoẻ Tâm thần ở Cấp độ Y tế Công cộng tại các Quốc gia Thu nhập Thấp và Trung bình: Nghiên cứu Hỗn hợp Thăm dò tại Việt Nam." PLOS One.'),
    ('Hoang M.D., B. Weiss, T. Lam and H. Ho',
     'Hoàng Minh Đặng, Weiss B., Lâm T. và Hồ H. 2018. "Kiến thức Sức khoẻ Tâm thần và Thích ứng Chương trình Can thiệp trong Quốc tế hoá Tâm lý học Trường học cho Việt Nam." Tâm lý học trong Trường học 55(8):941–54.'),
    ('Jones E.A., A.K. Mitra',
     'Jones E.A., Mitra A.K. và Bhuiyan A.R. 2021. "Tác động của COVID-19 lên Sức khoẻ Tâm thần ở Vị thành niên: Tổng quan Hệ thống." Tạp chí Quốc tế về Nghiên cứu Môi trường và Y tế Công cộng 18(5):2470.'),
    ('Kamimura, A. et al. 2018',
     'Kamimura A. và cộng sự. 2018. "Nhận thức về Sức khoẻ Tâm thần và Dịch vụ Sức khoẻ Tâm thần ở Sinh viên Đại học tại Việt Nam và Hoa Kỳ." Tạp chí Tâm thần học Châu Á 37:15–19.'),
    ('Keynejad, R.C., J. Spagnolo',
     'Keynejad R.C., Spagnolo J. và Thornicroft G. 2022. "Chăm sóc Sức khoẻ Tâm thần tại Bối cảnh Y tế Ban đầu và Cộng đồng: Bằng chứng Vượt ra ngoài Hướng dẫn Can thiệp mhGAP của WHO." Sức khoẻ Tâm thần Dựa trên Bằng chứng.'),
    ('Kim, J.H.J., W. Tsai',
     'Kim J.H.J., Tsai W., Kodish T., Lê T. Trung, Lau A.S. và Weiss B. 2019. "Khác biệt Văn hoá trong Mối liên hệ Thời gian giữa Khiếu nại Cơ thể, Triệu chứng Lo âu và Trầm cảm ở Vị thành niên." Tạp chí Nghiên cứu Tâm lý học Cơ thể 124:109763.'),
    ('Mai D., N.N.K. Pham',
     'Mai Đỗ, Phạm N.N.K., Wallick S. và Nastasi B.K. 2014. "Nhận thức về Bệnh Tâm thần và Kỳ thị Liên quan trong Quần thể người Việt: Phát hiện từ Nghiên cứu Phương pháp Hỗn hợp." Tạp chí Sức khoẻ Người nhập cư và Dân tộc Thiểu số 16:1294–98.'),
    ('Mckelvey, R.S., L.V. Baldassar',
     'McKelvey R.S., Baldassar L.V., Đỗ L. Sang và Roberts L. 1999. "Nhận thức của Phụ huynh Việt Nam về Bệnh Tâm thần ở Trẻ em và Vị thành niên." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 38(10):1302–09.'),
    ('Nga L.L., I. Shochet',
     'La Linh Nga, Shochet I., Trần T., Fisher J., Wurfl A., Nguyễn N., Orr J., Stocker R. và Nguyễn H. 2022. "Thích ứng Chương trình Sức khoẻ Tâm thần Trường học cho Vị thành niên tại Việt Nam." PLOS One 17(8): e0271959.'),
    ('Nguyen, T., T. Tran, H. Tran',
     'Nguyễn T., Trần T., Trần H., Trần T. và Fisher J. 2019. "Thách thức trong Lồng ghép Sức khoẻ Tâm thần vào Chăm sóc Y tế Ban đầu tại Việt Nam." trong Đổi mới Sức khoẻ Tâm thần Toàn cầu, do S. Okpaku biên tập: Springer, Cham.'),
    ('Nguyen, T.Q.C and T.H. Nguyen',
     'Nguyễn T.Q.C và Nguyễn T.H. 2018. "Kiến thức Sức khoẻ Tâm thần: Hiểu biết về Trầm cảm ở Sinh viên Đại học tại Hà Nội, Việt Nam." Tạp chí Quốc tế về Hệ thống Sức khoẻ Tâm thần 24(12):19.'),
    ('Ormel, J. et al. 2017',
     'Ormel J. và cộng sự. 2017. "Kết quả Chức năng của Rối loạn Tâm thần ở Trẻ em và Vị thành niên. Rối loạn Hiện tại Quan trọng nhất nhưng Lịch sử Tâm thần cũng Quan trọng." Y học Tâm lý 47(7):1271–82.'),
    ('Pagliaro, C., M. Pearl',
     'Pagliaro C., Pearl M., Lawrence D., Scott J.G. và Diminic S. 2021. "Ước tính Nhu cầu Chăm sóc Sức khoẻ Tâm thần ở Trẻ em và Vị thành niên Úc: Phát hiện từ Khảo sát Young Minds Matter." Tạp chí Tâm thần học Úc và New Zealand 56(11):1443–54.'),
    ('Patton G.C. et al. 2016',
     'Patton G.C. và cộng sự. 2016. "Tương lai của Chúng ta: Uỷ ban The Lancet về Sức khoẻ và Hạnh phúc Vị thành niên." The Lancet 387(10036):2423–78.'),
    ('Person, S., C. Hagquist',
     'Person S., Hagquist C. và Michelson D. 2017. "Tiếng nói Trẻ trong Chăm sóc Sức khoẻ Tâm thần: Khám phá Trải nghiệm và Ưu tiên Dịch vụ của Trẻ em và Vị thành niên." Tâm lý học và Tâm thần học Trẻ em Lâm sàng 22(1):140–51.'),
    ('Samuels, F., N. Jones, T. Gupta',
     'Samuels F., Jones N., Gupta T., Đặng B.T. và Lê D.H. 2018. Sức khoẻ Tâm thần và Hạnh phúc Tâm lý Xã hội ở Trẻ em và Người trẻ tại các Tỉnh và Thành phố được Chọn tại Việt Nam. Hà Nội: UNICEF.'),
    ('Schnyder, N., D. Lawrence',
     'Schnyder N., Lawrence D., Panczak R., Sawyer M.G., Whiteford H.A., Burgess P.M. và Harris M.G. 2019. "Nhu cầu được cảm nhận và rào cản đối với chăm sóc sức khoẻ tâm thần vị thành niên: sự đồng ý giữa vị thành niên và phụ huynh." Dịch tễ học và Khoa học Tâm thần 29(e60).'),
    ('Shaffer, D., P. Fisher, C.P. Lucas',
     'Shaffer D., Fisher P., Lucas C.P., Dulcan M.K. và Stone M.E.S. 2000. "Bảng Phỏng vấn Chẩn đoán NIMH cho Trẻ em Phiên bản IV (NIMH DISC-IV): Mô tả, Khác biệt so với Phiên bản Trước, và Độ tin cậy của một số Chẩn đoán Phổ biến." Tạp chí Viện Hàn lâm Tâm thần Trẻ em và Vị thành niên Hoa Kỳ 39(1):28–38.'),
    ('Tran, T., H.T. Nguyen, I. Shochet',
     'Trần T., Nguyễn H.T., Shochet I., Wurfl A., Orr J., Nguyễn N., La N., Nguyễn H., Stocker R., Nguyễn T., Lê M. và Fisher J. 2020. "Thử nghiệm Đối chứng Song song Hai nhánh tại Trường học về Can thiệp Bền bỉ Thích ứng Văn hoá để Cải thiện Sức khoẻ Tâm thần Vị thành niên tại Việt Nam: Đề cương Nghiên cứu." BMJ Open 10(10).'),
    ('Truc, T.T., N.L.L.T. Vu',
     'Thái T. Trúc, Vũ N.L.L.T. và Bùi H.H.T. 2020. "Kiến thức Sức khoẻ Tâm thần và Ưu tiên Tìm kiếm Trợ giúp ở Học sinh Trung học Phổ thông tại TP Hồ Chí Minh, Việt Nam." Sức khoẻ Tâm thần Trường học 12:378–87.'),
    ('UNICEF. 2020a',
     'UNICEF. 2020a. "Đánh giá Nhanh Tác động Xã hội và Kinh tế của COVID-19 lên Trẻ em và Gia đình tại Việt Nam." Hà Nội: UNICEF.'),
    ('UNICEF. 2020b',
     'UNICEF. 2020b. "Tác động của COVID-19 lên Sức khoẻ Tâm thần của Vị thành niên và Thanh niên."'),
    ('United Nations Children',
     'Quỹ Nhi đồng Liên Hợp Quốc (UNICEF) và Viện Phát triển Hải ngoại (ODI). 2018. "Bản chất của Tự sát ở Trẻ em và Người trẻ tại các Tỉnh và Thành phố được Chọn tại Việt Nam." Hà Nội: UNICEF.'),
    ('Wasserman, D., H.T.T. Tran',
     'Wasserman D., Trần H.T.T., Phạm D.T.M., Goldstein M., Nordenskiöld A. và Wasserman C. 2008. "Quá trình Tự sát, Giao tiếp về Tự sát và Tình huống Tâm lý Xã hội của Người trẻ Toan Tự sát trong một Cộng đồng Nông thôn Việt Nam." Tâm thần học Thế giới 7(1):47–53.'),
    ('Weiss, B., M. Dang, L. Trung',
     'Weiss B., Đặng M., Lê Trung, Nguyễn M.C., Nguyễn H.T.T. và Pollack A. 2014. "Đánh giá Dịch tễ học và Yếu tố Nguy cơ Đại diện Quốc gia về Sức khoẻ Tâm thần Trẻ em tại Việt Nam." International Perspectives in Psychology 3(3):139–53.'),
    ('World Health Organization. 2014',
     'Tổ chức Y tế Thế giới. 2014. "Sức khoẻ cho Vị thành niên Thế giới: Cơ hội Thứ Hai trong Thập kỷ Thứ Hai." Geneva.'),
    ('World Health Organization. 2016',
     'Tổ chức Y tế Thế giới. 2016. Hướng dẫn Can thiệp mhGAP cho Rối loạn Tâm thần, Thần kinh và Sử dụng Chất trong các Cơ sở Y tế Không Chuyên khoa: Chương trình Hành động Lấp khoảng trống Sức khoẻ Tâm thần (mhGAP) – Phiên bản 2.0. Geneva.'),
    ('World Health Organization, Ministry of Health',
     'Tổ chức Y tế Thế giới, Bộ Y tế và Bộ Giáo dục & Đào tạo. 2022. Khảo sát Sức khoẻ Học sinh Toàn cầu tại Trường học tại Việt Nam 2019: Báo cáo. Manila: Văn phòng Khu vực WHO Tây Thái Bình Dương.'),
]

d = Document(OUT)

# Identify refs section (start after "TÀI LIỆU THAM KHẢO" + "(Theo chuẩn")
refs_start = None
for i, p in enumerate(d.paragraphs):
    if p.text.startswith('American Psychiatric Association. 2013'):
        refs_start = i
        break
print(f'Refs start at paragraph #{refs_start}')

added = 0
for prefix, vn_text in BILING:
    # Search only in refs section
    for i in range(refs_start or 0, len(d.paragraphs)):
        p = d.paragraphs[i]
        if p.text.startswith(prefix):
            # Check if next paragraph already has a VN translation (don't double-add)
            nxt_idx = i + 1
            # Skip empty paragraphs
            while nxt_idx < len(d.paragraphs) and not d.paragraphs[nxt_idx].text.strip():
                nxt_idx += 1
            nxt = d.paragraphs[nxt_idx].text if nxt_idx < len(d.paragraphs) else ''
            if nxt.startswith(vn_text[:40]) or (vn_text[:25] in nxt):
                # Already added
                break
            # Add Vietnamese paragraph immediately after current
            new_para = d.add_paragraph()
            r = new_para.add_run(vn_text)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(10)
            r.italic = True
            r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
            en_el = p._element
            new_el = new_para._element
            new_el.getparent().remove(new_el)
            en_el.addnext(new_el)
            added += 1
            print(f'  + {prefix[:45]:48s}')
            break

print(f'\nTotal EN refs translated: {added}/{len(BILING)}')
d.save(OUT)
