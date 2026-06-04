# -*- coding: utf-8 -*-
"""Sinh Trich yeu Luan an cho NCS Cong Thi Hang.
- Theo format template Tran Anh Khoi (A4, TNR 13pt body, margins 2.54cm all)
- Voi VN: Cong Thi Hang voice + content tu LA chinh + Ket luan
- Voi EN: dich literal-faithful
29/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', 'TrichYeuLA_CongThiHang_29052026.docx')

d = Document()

# ============================================================
# PAGE SETUP - khop template
# ============================================================
for sec in d.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

# Style Normal
s = d.styles['Normal']
s.font.name = 'Times New Roman'
s.font.size = Pt(13)
s.paragraph_format.line_spacing = 1.15
s.paragraph_format.space_after = Pt(6)


# ============================================================
# HELPERS
# ============================================================
def H(text, size=13, bold=True, italic=False, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=6):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def F(label, value, label_bold=True, value_bold=True, size=13):
    """Field: label + value, both on same paragraph."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_after = Pt(4)
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(size)
    r1.bold = label_bold
    r2 = p.add_run(value)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(size)
    r2.bold = value_bold
    return p


def P(text, size=13, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
      space_after=6, first_line=0.75):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if first_line is not None:
        p.paragraph_format.first_line_indent = Cm(first_line)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_signature_table(left_lines, right_lines):
    """Bang chu ky 2 cot."""
    t = d.add_table(rows=1, cols=2)
    t.autofit = True
    cell_left = t.rows[0].cells[0]
    cell_right = t.rows[0].cells[1]

    # Clear default paragraph
    cell_left.text = ''
    cell_right.text = ''

    for i, line in enumerate(left_lines):
        p = cell_left.paragraphs[0] if i == 0 else cell_left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        # Bold for first line and last line (role + name)
        if i == 0 or i == len(left_lines) - 1:
            r.bold = True

    for i, line in enumerate(right_lines):
        p = cell_right.paragraphs[0] if i == 0 else cell_right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        if i == 0 or i == len(right_lines) - 1:
            r.bold = True
    return t


# ============================================================
# VIETNAMESE SECTION
# ============================================================

# Title
H('TRÍCH YẾU LUẬN ÁN', size=13, bold=True, space_after=12)

# Empty line
d.add_paragraph()

# Fields
F('Tên đề tài: ', 'Rối loạn lo âu ở học sinh trung học cơ sở')
F('Chuyên ngành: ', 'Tâm lý học chuyên ngành')
F('Mã số: ', '9310401')
F('Nghiên cứu sinh: ', '\t\tCông Thị Hằng')
F('Cán bộ hướng dẫn: ', '\t\tTS. Đào Minh Đức')
F('Cơ sở đào tạo: ', 'Trường Đại học Sư phạm Hà Nội')

d.add_paragraph()

# Heading: Nhung ket luan moi
H('Những kết luận mới của luận án', size=13, bold=True, space_after=10)

# Dong gop ly luan
P_lyluan_heading = d.add_paragraph()
P_lyluan_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
P_lyluan_heading.paragraph_format.space_after = Pt(6)
r_ll = P_lyluan_heading.add_run('Đóng góp về mặt lý luận')
r_ll.font.name = 'Times New Roman'
r_ll.font.size = Pt(13)
r_ll.bold = True

P('Luận án đã hệ thống hóa và làm rõ cơ sở lý luận về rối loạn lo âu ở học sinh '
  'trung học cơ sở theo hướng tiếp cận đa chiều, bao gồm ba dạng rối loạn lo âu '
  'phổ biến nhất ở trẻ vị thành niên theo phân loại DSM-5 (rối loạn lo âu lan tỏa, '
  'rối loạn lo âu xã hội và rối loạn lo âu chia ly). Đồng thời, nghiên cứu xác lập '
  'rối loạn lo âu như một hệ thống đa thành phần, bao gồm các khía cạnh nhận thức, '
  'cảm xúc, sinh lý và hành vi, qua đó làm rõ đặc điểm tâm lý đặc thù của lứa tuổi '
  'học sinh trung học cơ sở trong giai đoạn phát triển dậy thì — giai đoạn có nhiều '
  'biến đổi về nhận thức bản thân và quan hệ xã hội.')

P('Bên cạnh đó, luận án xác lập mô hình mối quan hệ giữa rối loạn lo âu với hệ thống '
  'các yếu tố ảnh hưởng, bao gồm ba nhóm yếu tố nguy cơ (áp lực học tập, nghiện điện '
  'thoại, bắt nạt học đường), bốn nhóm yếu tố bảo vệ (lòng tự trọng, gắn bó trường học, '
  'hỗ trợ của cha mẹ, hỗ trợ của bạn bè) và biện pháp đối phó của học sinh. Đóng góp '
  'này góp phần lấp đầy khoảng trống nghiên cứu về phương pháp luận trong lĩnh vực sức '
  'khỏe tâm thần học đường ở Việt Nam, đặc biệt ở quy trình kiểm định chất lượng công '
  'cụ đo lường và chuẩn hóa thang đo phù hợp với ngôn ngữ, văn hóa Việt Nam.')

# Dong gop thuc tien
P_thuctien_heading = d.add_paragraph()
P_thuctien_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
P_thuctien_heading.paragraph_format.space_after = Pt(6)
r_tt = P_thuctien_heading.add_run('Đóng góp về mặt thực tiễn')
r_tt.font.name = 'Times New Roman'
r_tt.font.size = Pt(13)
r_tt.bold = True

P('Về phương pháp đo lường, luận án đã hiệu chỉnh và kiểm định độ tin cậy, độ giá trị '
  'của tám thang đo tâm lý (RCADS, ESSA, MPAS, MPVS, PSSM, MSPSS, Brief COPE và '
  'Rosenberg) thông qua các chỉ số Cronbach\'s alpha, McDonald\'s omega và phân tích '
  'nhân tố khẳng định (CFA). Các thang đo đều đạt ngưỡng chấp nhận, tạo ra một bộ '
  'công cụ đo lường có độ ổn định và độ tin cậy cao, phù hợp với bối cảnh học sinh '
  'trung học cơ sở tại Việt Nam và là tiền đề cho các nghiên cứu tiếp theo trong '
  'lĩnh vực sức khỏe tâm thần học đường.')

P('Kết quả nghiên cứu thực trạng trên mẫu 1.352 học sinh trung học cơ sở tại Hà Nội đã '
  'kiểm chứng thành công ba giả thuyết khoa học: (1) học sinh nữ có mức độ rối loạn '
  'lo âu lan tỏa và rối loạn lo âu xã hội cao hơn học sinh nam, trong khi rối loạn lo '
  'âu chia ly không có khác biệt theo giới tính (F = 0,246; p > 0,05); (2) lòng tự '
  'trọng là yếu tố bảo vệ có cường độ tác động mạnh, tương đương khoảng 86% so với áp '
  'lực học tập trong mô hình phương trình cấu trúc (SEM); (3) ba nhóm yếu tố nguy cơ '
  'có thứ tự cường độ tác động từ cao xuống thấp là: áp lực học tập > nghiện điện '
  'thoại > bắt nạt thể chất. Bên cạnh đó, nghiên cứu phát hiện các chiến lược đối phó '
  'hiện tại của học sinh trung học cơ sở chưa thực sự hiệu quả, thể hiện qua xu hướng '
  'gia tăng triệu chứng rối loạn lo âu khi mức độ sử dụng các biện pháp đối phó tăng — '
  'cho thấy cần định hướng học sinh chuyển từ phản ứng né tránh sang các chiến lược '
  'chủ động như giải quyết vấn đề và tìm kiếm hỗ trợ xã hội.')

P('Trên cơ sở các phát hiện thực nghiệm, luận án đề xuất Khung chương trình tập huấn '
  'phòng ngừa rối loạn lo âu cho học sinh trung học cơ sở gồm 8 nội dung cơ bản, đáp '
  'ứng yêu cầu cấp thiết hiện nay về chăm sóc sức khỏe tâm thần học đường ở Việt Nam. '
  'Khung chương trình này cung cấp cơ sở khoa học cho các kiến nghị can thiệp đối với '
  'học sinh, giáo viên, gia đình và các lực lượng hỗ trợ học đường (chuyên viên tâm '
  'lý, phòng tư vấn học đường, đoàn thanh niên, các câu lạc bộ), đồng thời là tài liệu '
  'tham khảo có giá trị cho giáo viên, chuyên viên tâm lý học đường và nhà quản lý '
  'giáo dục trong công tác phòng ngừa và hỗ trợ học sinh có biểu hiện rối loạn lo âu.')

d.add_paragraph()

# Signature table VN
add_signature_table(
    left_lines=['Người hướng dẫn', '', '', '', 'TS. Đào Minh Đức'],
    right_lines=['Nghiên cứu sinh', '', '', '', 'Công Thị Hằng'],
)

# Page break
d.add_page_break()


# ============================================================
# ENGLISH SECTION
# ============================================================

H('DISSERTATION ABSTRACT', size=13, bold=True, space_after=12)
d.add_paragraph()

F('Title: ', 'Anxiety Disorders among Lower Secondary School Students')
F('Speciality: ', 'Psychology')
F('Classification: ', '9310401')
F('Name of PhD Student: ', '\t\tCong Thi Hang')
F('Advisor: ', '\t\tDr. Dao Minh Duc')
F('Institutional: ', 'Hanoi National University of Education')

d.add_paragraph()

H('New conclusions', size=13, bold=True, space_after=10)

# Theoretical contribution
P_th_heading = d.add_paragraph()
P_th_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
P_th_heading.paragraph_format.space_after = Pt(6)
r_th = P_th_heading.add_run('Theoretical contribution')
r_th.font.name = 'Times New Roman'
r_th.font.size = Pt(13)
r_th.bold = True
r_th.italic = True

P('The dissertation has systematized and clarified the theoretical foundations of '
  'anxiety disorders among lower secondary school students through a multidimensional '
  'approach, encompassing the three most prevalent anxiety disorder types in adolescents '
  'according to the DSM-5 classification: generalized anxiety disorder, social anxiety '
  'disorder, and separation anxiety disorder. Furthermore, the study conceptualizes '
  'anxiety disorders as a multi-component system comprising cognitive, emotional, '
  'physiological, and behavioral dimensions, thereby illuminating the distinctive '
  'psychological characteristics of lower secondary school students during the pubertal '
  'developmental stage — a period marked by substantial transformations in self-perception '
  'and social relationships.')

P('In addition, the dissertation establishes a model of the relationships between '
  'anxiety disorders and a system of influencing factors, including three risk factor '
  'groups (academic pressure, smartphone addiction, school bullying), four protective '
  'factor groups (self-esteem, school engagement, parental support, peer support), and '
  'students\' coping strategies. This contribution helps fill the methodological gap in '
  'the field of school-based mental health research in Vietnam, particularly regarding '
  'the validation procedures for measurement instruments and the standardization of '
  'psychological scales to align with the Vietnamese language and cultural context.')

# Practical contributions
P_pr_heading = d.add_paragraph()
P_pr_heading.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
P_pr_heading.paragraph_format.space_after = Pt(6)
r_pr = P_pr_heading.add_run('Practical contributions')
r_pr.font.name = 'Times New Roman'
r_pr.font.size = Pt(13)
r_pr.bold = True
r_pr.italic = True

P('Methodologically, the dissertation has adapted and validated the reliability and '
  'validity of eight psychological scales (RCADS, ESSA, MPAS, MPVS, PSSM, MSPSS, Brief '
  'COPE, and Rosenberg) through Cronbach\'s alpha, McDonald\'s omega, and confirmatory '
  'factor analysis (CFA). All scales attained acceptable thresholds, thereby '
  'establishing a measurement toolkit with high stability and reliability suitable for '
  'the lower secondary school context in Vietnam and laying the groundwork for '
  'subsequent research in the field of school-based mental health.')

P('The empirical findings from a sample of 1,352 lower secondary school students in Hanoi '
  'successfully confirmed three scientific hypotheses: (1) female students exhibit higher '
  'levels of generalized anxiety disorder and social anxiety disorder than male students, '
  'whereas separation anxiety shows no significant gender differences (F = 0.246; '
  'p > 0.05); (2) self-esteem serves as a protective factor with a strong magnitude, '
  'approximately 86% comparable to academic pressure within the structural equation '
  'model (SEM); (3) the three risk factor groups, in descending order of magnitude, are: '
  'academic pressure > smartphone addiction > physical bullying. Additionally, the '
  'study reveals that the current coping strategies of lower secondary school students '
  'are not yet sufficiently effective, evidenced by the upward trend in anxiety symptoms '
  'accompanying increased use of coping behaviors — indicating the need to guide '
  'students from avoidance reactions toward proactive strategies such as problem-solving '
  'and social support seeking.')

P('Building upon these empirical findings, the dissertation proposes an Eight-Component '
  'Framework for the Prevention of Anxiety Disorders among lower secondary school '
  'students, addressing the urgent contemporary need for school-based mental health '
  'care in Vietnam. This framework provides a scientific basis for intervention '
  'recommendations targeting students, teachers, families, and school support personnel '
  '(school psychologists, school counseling offices, youth unions, student clubs), '
  'while also serving as a valuable reference document for teachers, school '
  'psychologists, and education administrators in the prevention of and support for '
  'students manifesting anxiety symptoms.')

d.add_paragraph()

# Signature table EN
add_signature_table(
    left_lines=['Advisor', '', '', '', 'Dr. Dao Minh Duc'],
    right_lines=['PhD Candidate', '', '', '', 'Cong Thi Hang'],
)


# ============================================================
# CLEAN METADATA
# ============================================================
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
cp.identifier = ''
cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'File size: {os.path.getsize(OUT)} bytes')
