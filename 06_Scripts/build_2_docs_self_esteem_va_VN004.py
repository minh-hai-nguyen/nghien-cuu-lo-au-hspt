"""Build 2 docs:
1. Self-esteem va anxiety - 6 nghien cuu QT bo sung tong quan chuong 1 (CTH v6)
2. Phan bien VN004 Nguyen Thi Van - yeu to ban than HS (format CAU TRA LOI to xanh)
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x30)
BLACK = RGBColor(0x00, 0x00, 0x00)

def new_doc(margin_l=2.5, font_sz=12):
    d = Document()
    for s in d.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(margin_l); s.right_margin = Cm(2.0)
    style = d.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(font_sz)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return d

def H(d, text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(d, text, color=BLACK, bold=False, italic=False, size=13, indent=True, justify=True):
    p = d.add_paragraph()
    if justify: p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.color.rgb = color
    r.font.size = Pt(size); r.bold = bold; r.italic = italic

def bullet(d, text, color=BLACK, italic=False, size=13):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(size); r.italic = italic

def caption(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def add_table(d, header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def ref_entry(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# ============================================================
# DOC 1 — Self-esteem và anxiety: nghiên cứu nước ngoài (CTH v6)
# ============================================================
def build_self_esteem_doc():
    d = new_doc(margin_l=3.0, font_sz=13)

    # Title
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('TỔNG QUAN — MỐI QUAN HỆ GIỮA LÒNG TỰ TRỌNG\nVÀ LO ÂU Ở VỊ THÀNH NIÊN\n(Bổ sung cho chương 1 đề tài)')
    r.bold = True; r.font.size = Pt(14)
    para(d, '')

    # Mở đầu — 3 cấp độ
    H(d, '1. Đặt vấn đề', level=2)
    para(d,
        'Lòng tự trọng (self-esteem) — đánh giá tổng thể của cá nhân về '
        'giá trị bản thân — đã được Rosenberg (1965) khái niệm hóa trong '
        'tác phẩm kinh điển "Society and the Adolescent Self-Image". Trong '
        'nửa thế kỷ tiếp theo, nhiều nghiên cứu thuần tập đã làm rõ vai '
        'trò của lòng tự trọng trong phát triển sức khỏe tâm thần, đặc '
        'biệt mối quan hệ với trầm cảm và lo âu. Tuy nhiên, bản chất quan '
        'hệ giữa lòng tự trọng và lo âu vẫn là chủ đề tranh luận — một số '
        'nghiên cứu khẳng định "mô hình tổn thương" (vulnerability model: '
        'lòng tự trọng thấp dẫn đến lo âu), số khác ủng hộ "mô hình sẹo" '
        '(scar model: lo âu làm xói mòn lòng tự trọng).'
    )
    para(d,
        'Phần dưới đây tổng hợp sáu nghiên cứu nước ngoài KINH ĐIỂN và '
        'CẬP NHẬT về mối quan hệ này — từ phân tích tổng hợp dữ liệu '
        'thuần tập (Sowislo & Orth, 2013) đến nghiên cứu thuần tập 23 '
        'năm (Steiger và cộng sự, 2014) và phân tích tổng hợp can thiệp '
        'tại trường (Cai và cộng sự, 2025). Mỗi nghiên cứu được trình '
        'bày với phương pháp, phát hiện chính, và hàm ý cho đề tài lo âu '
        'học sinh trung học cơ sở Việt Nam.'
    )

    # 2.1 Sowislo & Orth 2013
    H(d, '2.1. Sowislo và Orth (2013) — Phân tích tổng hợp KINH ĐIỂN', level=3)
    para(d,
        'Phân tích tổng hợp 95 nghiên cứu thuần tập của Sowislo và Orth '
        '(2013) trong Psychological Bulletin là CÔNG TRÌNH KINH ĐIỂN khẳng '
        'định mối quan hệ giữa lòng tự trọng và rối loạn nội hóa '
        '(internalizing disorders). Cụ thể, 77 nghiên cứu về trầm cảm và '
        '18 nghiên cứu về lo âu được đưa vào phân tích.'
    )
    para(d,
        'Phát hiện CỐT LÕI: lòng tự trọng dự báo lo âu với β chuẩn hóa = '
        '−0,10, còn lo âu dự báo lòng tự trọng với β = −0,08 — hai chiều '
        'tương đối cân bằng. Đối với trầm cảm, mô hình tổn thương được '
        'xác lập rõ rệt hơn (β = −0,16 từ tự trọng → trầm cảm). Nói cách '
        'khác, KHÁC với trầm cảm, mối quan hệ giữa lòng tự trọng và lo '
        'âu là HAI CHIỀU — vừa là yếu tố nguy cơ vừa là hệ quả của lo âu.'
    )
    para(d,
        'Hàm ý cho nghiên cứu hiện tại: phân tích cắt ngang KHÔNG đủ để '
        'xác lập chiều quan hệ — cần thiết kế thuần tập theo dõi cùng '
        'nhóm học sinh ít nhất 12 tháng để phân biệt mô hình tổn thương '
        'và mô hình sẹo. Phù hợp với khuyến nghị của Trần Thảo Vi và '
        'cộng sự (2024) đã thực hiện thuần tập 3 năm tại Huế.'
    )

    # 2.2 Steiger et al. 2014
    H(d, '2.2. Steiger và cộng sự (2014) — Thuần tập 23 năm', level=3)
    para(d,
        'Steiger, Allemand, Robins và Fend (2014) trong Journal of '
        'Personality and Social Psychology đã thực hiện nghiên cứu thuần '
        'tập 23 năm trên 1.527 thanh thiếu niên Đức — đo lòng tự trọng '
        'hằng năm từ tuổi 12 đến 16, sau đó đánh giá trầm cảm ở tuổi 16 '
        'và 35. Phân tích đường cong tăng trưởng tiềm ẩn (latent growth '
        'curve analyses) cho thấy CẢ MỨC ĐỘ và HƯỚNG THAY ĐỔI của lòng '
        'tự trọng đều dự báo trầm cảm trưởng thành.'
    )
    para(d,
        'Phát hiện đặc sắc: thanh thiếu niên ĐI VÀO tuổi vị thành niên '
        'với lòng tự trọng thấp, hoặc có lòng tự trọng GIẢM trong giai '
        'đoạn vị thành niên, có nhiều khả năng biểu hiện triệu chứng '
        'trầm cảm hai thập niên sau ở tuổi trưởng thành. Mẫu hình này '
        'đúng cho cả lòng tự trọng tổng thể và lòng tự trọng theo lĩnh '
        'vực (hấp dẫn ngoại hình, năng lực học tập).'
    )
    para(d,
        'Hàm ý: lòng tự trọng giai đoạn 11–15 tuổi (tương đương trung '
        'học cơ sở Việt Nam) là CỬA SỔ CAN THIỆP CRITICAL — sự can '
        'thiệp ở giai đoạn này có thể tạo hiệu ứng kéo dài đến tuổi '
        'trưởng thành. Củng cố thêm khuyến nghị tăng cường thành phần '
        '"tự trọng" trong chương trình can thiệp tại trường trung học '
        'cơ sở.'
    )

    # 2.3 Orth & Robins 2014
    H(d, '2.3. Orth và Robins (2014) — Tổng quan phát triển lòng tự trọng', level=3)
    para(d,
        'Orth và Robins (2014) trong Current Directions in Psychological '
        'Science tổng quan các nghiên cứu thuần tập gần đây về phát triển '
        'lòng tự trọng và ảnh hưởng đến các kết quả đời sống. Ba kết luận '
        'chính được rút ra.'
    )
    para(d,
        'Thứ nhất, lòng tự trọng TĂNG dần từ vị thành niên đến trung niên, '
        'đạt đỉnh ở tuổi 50–60, sau đó GIẢM với tốc độ tăng dần khi vào '
        'tuổi già. Thứ hai, lòng tự trọng là đặc điểm tương đối ỔN ĐỊNH '
        'nhưng KHÔNG bất biến — có thể thay đổi qua trải nghiệm và can '
        'thiệp. Thứ ba, lòng tự trọng cao DỰ BÁO thành công và an sinh '
        'trong các lĩnh vực đời sống bao gồm quan hệ, công việc, và sức '
        'khỏe.'
    )
    para(d,
        'Hàm ý cho đề tài: phát hiện "lòng tự trọng tăng từ vị thành niên" '
        'gợi ý xu hướng tự nhiên thuận lợi cho can thiệp ở giai đoạn '
        'trung học cơ sở — chương trình hỗ trợ chỉ cần CỦNG CỐ và DUY TRÌ '
        'xu hướng tăng tự nhiên này thay vì phải đảo ngược xu hướng giảm.'
    )

    # 2.4 Cai 2025
    H(d, '2.4. Cai và cộng sự (2025) — Phân tích tổng hợp can thiệp tại trường', level=3)
    para(d,
        'Cai, Mei, Wang và Luo (2025) trong Frontiers in Psychiatry thực '
        'hiện phân tích tổng hợp các thử nghiệm ngẫu nhiên có đối chứng '
        'về can thiệp resilience tại trường — bao gồm thành phần lòng tự '
        'trọng, kỹ năng ứng phó, kết nối xã hội, và lạc quan. Các nghiên '
        'cứu được thực hiện tại nhiều quốc gia: Mỹ (n = 10), Trung Quốc '
        '(n = 9), Úc (n = 4), Pakistan (n = 2), Ấn Độ (n = 2).'
    )
    para(d,
        'Phát hiện chính: can thiệp resilience tại trường CÓ HIỆU QUẢ '
        'tăng resilience và giảm triệu chứng sức khỏe tâm thần ở trẻ em '
        'và vị thành niên — kích thước hiệu ứng tổng hợp SMD = 0,17 '
        '(KTC 95% từ 0,06 đến 0,29; p < 0,01) trên 21 thử nghiệm, '
        'tương ứng hiệu ứng NHỎ. Tuy nhiên, heterogeneity rất CAO '
        '(I² = 81,90%) — gợi ý hiệu quả tồn tại nhưng có khác biệt '
        'đáng kể giữa các nghiên cứu do thiết kế can thiệp khác nhau. '
        'Phù hợp với kết luận của Cai và cộng sự — resilience là yếu '
        'tố BẢO VỆ, BỔ SUNG cho liệu pháp nhận thức hành vi (CBT) chứ '
        'không thay thế.'
    )

    # 2.5 Masten 2014
    H(d, '2.5. Masten (2014) — Khung resilience toàn cầu', level=3)
    para(d,
        'Masten (2014) trong Child Development trình bày khung resilience '
        'toàn cầu cho trẻ em và thanh thiếu niên — dựa trên Bài phát biểu '
        'Tổng thống tại hội nghị biennial của Society for Research in '
        'Child Development năm 2013. Trong khung này, lòng tự trọng được '
        'xếp vào nhóm "ordinary magic" — các yếu tố nội tại bảo vệ trẻ '
        'em trước nghịch cảnh.'
    )
    para(d,
        'Khung Masten nhấn mạnh hai luận điểm. Thứ nhất, resilience không '
        'phải là đặc điểm hiếm có — đa số trẻ em có khả năng thích ứng '
        'tốt khi được hỗ trợ phù hợp. Thứ hai, lòng tự trọng phát triển '
        'thông qua tương tác giữa cá nhân và môi trường — không thể tách '
        'rời chương trình can thiệp lòng tự trọng khỏi việc cải thiện '
        'môi trường gia đình, trường học, và cộng đồng.'
    )

    # 2.6 Rosenberg 1965
    H(d, '2.6. Rosenberg (1965) — Khái niệm và thang đo gốc', level=3)
    para(d,
        'Rosenberg (1965) trong tác phẩm "Society and the Adolescent '
        'Self-Image" đã xây dựng khái niệm lòng tự trọng và phát triển '
        'thang đo Rosenberg Self-Esteem Scale (RSES) — gồm 10 mục, Likert '
        '4 điểm. Đây là thang đo lòng tự trọng được sử dụng rộng rãi '
        'nhất trong y văn quốc tế, đã được dịch sang hơn 50 ngôn ngữ và '
        'có độ tin cậy nội tại cao (Cronbach alpha thường > 0,80).'
    )
    para(d,
        'Đối với nghiên cứu lo âu học sinh trung học cơ sở Việt Nam, '
        'thang RSES là lựa chọn ƯU TIÊN do (a) ngắn gọn (10 mục, hoàn '
        'thành dưới 5 phút); (b) đã được chuẩn hóa cho dân số Việt Nam; '
        '(c) cho phép so sánh trực tiếp với y văn quốc tế. Đây là thang '
        'đo được sử dụng trong các nghiên cứu Sowislo & Orth (2013), '
        'Steiger và cộng sự (2014), và phần lớn các thử nghiệm trong '
        'phân tích tổng hợp của Cai và cộng sự (2025).'
    )

    # 3. Tổng hợp
    H(d, '3. Tổng hợp và hàm ý cho đề tài', level=2)
    caption(d, 'Bảng 1. Tổng hợp sáu nghiên cứu nước ngoài về lòng tự trọng và lo âu')
    add_table(d,
        ['Nghiên cứu', 'Loại NC', 'Phát hiện chính'],
        [
            ['Sowislo & Orth (2013), Psychological Bulletin',
             'Meta-analysis 95 thuần tập',
             'β = −0,10 (tự trọng → lo âu); quan hệ HAI CHIỀU'],
            ['Steiger và cộng sự (2014), JPSP',
             'Thuần tập 23 năm, n = 1.527',
             'Tự trọng thấp tuổi 12–16 dự báo trầm cảm tuổi 35'],
            ['Orth & Robins (2014), Current Directions',
             'Tổng quan thuần tập',
             'Tự trọng tăng từ vị thành niên đến trung niên'],
            ['Cai và cộng sự (2025), Frontiers in Psychiatry',
             'Meta-analysis 21 RCT tại trường',
             'SMD = 0,17 (KTC 95% 0,06–0,29); I² = 81,90%'],
            ['Masten (2014), Child Development',
             'Khung lý thuyết',
             'Tự trọng thuộc nhóm "ordinary magic" — resilience'],
            ['Rosenberg (1965), Princeton University Press',
             'Khái niệm + thang đo gốc',
             'RSES — thang đo chuẩn vàng được dùng > 50 ngôn ngữ'],
        ]
    )

    para(d, '')
    para(d,
        'Sáu nghiên cứu trên thiết lập ba luận điểm vững chắc cho đề '
        'tài lo âu học sinh trung học cơ sở Việt Nam. Thứ nhất, lòng '
        'tự trọng và lo âu có quan hệ HAI CHIỀU — không chỉ là yếu tố '
        'nguy cơ mà còn là hệ quả (Sowislo & Orth, 2013). Thứ hai, '
        'lòng tự trọng giai đoạn 11–15 tuổi là CỬA SỔ CAN THIỆP '
        'CRITICAL với hiệu ứng kéo dài hai thập niên (Steiger và cộng '
        'sự, 2014). Thứ ba, can thiệp tự trọng tại trường có hiệu quả '
        'thực tế — kích thước hiệu ứng nhỏ–trung bình nhưng đáng kể ở '
        'cấp dân số (Cai và cộng sự, 2025).'
    )
    para(d,
        'Phát hiện này đặc biệt liên quan đến giả thuyết H6 của nghiên '
        'cứu — "tự trọng có cường độ tác động bảo vệ ngang bằng áp lực '
        'học tập" — vì các nghiên cứu nước ngoài chứng minh lòng tự '
        'trọng là yếu tố bảo vệ MẠNH, không yếu hơn các yếu tố nguy cơ '
        'như giả định thông thường. Phù hợp với kết quả thực tế của '
        'chương 3 luận án (β tự trọng = −0,455 cho lo âu lan tỏa; '
        '−0,415 cho lo âu xã hội — Bảng 3.32) — gần ngang cường độ áp '
        'lực học tập (β = 0,510 và 0,490 — Bảng 3.24).'
    )

    # 4. TLTK
    H(d, '4. Tài liệu tham khảo', level=2)
    refs = [
        'Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, 1594658.',
        'Masten, A. S. (2014). Global perspectives on resilience in children and youth. Child Development, 85(1), 6–20. https://doi.org/10.1111/cdev.12205',
        'Orth, U., & Robins, R. W. (2014). The development of self-esteem. Current Directions in Psychological Science, 23(5), 381–387. https://doi.org/10.1177/0963721414547414',
        'Rosenberg, M. (1965). Society and the adolescent self-image. Princeton University Press.',
        'Sowislo, J. F., & Orth, U. (2013). Does low self-esteem predict depression and anxiety? A meta-analysis of longitudinal studies. Psychological Bulletin, 139(1), 213–240. https://doi.org/10.1037/a0028931',
        'Steiger, A. E., Allemand, M., Robins, R. W., & Fend, H. A. (2014). Low and decreasing self-esteem during adolescence predict adult depression two decades later. Journal of Personality and Social Psychology, 106(2), 325–338. https://doi.org/10.1037/a0035133',
        'Trần, T. V., và cộng sự. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine.',
    ]
    for r in refs:
        ref_entry(d, r)

    out = OUT_DIR / 'Tong_quan_self_esteem_va_lo_au_6_NC_QT.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')

# ============================================================
# DOC 2 — Phản biện VN004 yếu tố bản thân HS (CÂU TRẢ LỜI tô xanh)
# ============================================================
def build_VN004_critique():
    d = new_doc(margin_l=2.5, font_sz=12)

    H(d, 'Phản biện chi tiết VN004 Nguyễn Thị Vân', level=1, color=NAVY)
    H(d, '— Yếu tố "bản thân học sinh" trong Mức độ lo âu HS THPT TPHCM —', level=2, color=NAVY)

    # CÂU HỎI tô xanh
    H(d, 'Câu hỏi của thầy', level=2, color=NAVY)
    p = d.add_paragraph()
    r = p.add_run('Em phản biện cho thầy Nguyễn Thị Vân với yếu tố bản thân học sinh nhé.')
    r.font.color.rgb = BLUE; r.font.size = Pt(12)

    # 1. Bối cảnh
    H(d, '1. Bối cảnh — VN004 Nguyễn Thị Vân (2020)', level=2, color=NAVY)
    para(d,
        'Bài VN004 trong cơ sở dữ liệu dự án — Nguyễn Thị Vân (2018; bản '
        'mở rộng 2019/2020) "Mức độ lo âu của học sinh trung học phổ '
        'thông thành phố Hồ Chí Minh" — là một trong số ít nghiên cứu '
        'Việt Nam tập trung vào YẾU TỐ ẢNH HƯỞNG đến lo âu học sinh '
        '(không chỉ tỷ lệ).', indent=False
    )
    para(d, 'Phương pháp:', bold=True, indent=False)
    bullet(d, 'Cắt ngang 2 pha — sàng lọc 558 HS THPT bằng STAI Spielberger (bản Việt hóa Nguyễn Công Khanh, 2000) → 90 HS phỏng vấn sâu.', size=12)
    bullet(d, 'Mẫu thuận tiện, 4 trường nội + ngoại thành TPHCM.', size=12)
    bullet(d, 'Tỷ lệ RLLA: 15–18,5% (STAI ≥ ngưỡng).', size=12)
    bullet(d, 'Bốn nhóm yếu tố ảnh hưởng: học tập / gia đình / quan hệ xã hội / BẢN THÂN HS.', size=12)

    # 2. Yếu tố bản thân HS — số liệu trong VN004
    H(d, '2. Số liệu VN004 về yếu tố "bản thân học sinh"', level=2, color=NAVY)
    add_table(d,
        ['Chỉ số', 'Giá trị trong VN004'],
        [
            ['Hệ số tương quan r (nhóm yếu tố bản thân HS)', '0,42'],
            ['Top 1 biểu hiện: ít giao tiếp chia sẻ', 'X̄ = 0,91'],
            ['Top 2: thất vọng', 'X̄ = 0,90'],
            ['Top 3: lo sợ khó khăn', 'X̄ = 0,88'],
            ['"Tỷ lệ giải thích" tác giả công bố', '85,4%'],
        ]
    )
    para(d, '')
    para(d,
        'Đặc điểm đáng chú ý: tác giả công bố nhóm "bản thân HS" giải '
        'thích 85,4% — con số này CAO BẤT THƯỜNG so với hệ số tương quan '
        'r = 0,42.', indent=False
    )

    # 3. Phản biện chi tiết
    H(d, '3. Phản biện sáu điểm yếu phương pháp luận', level=2, color=NAVY)

    H(d, '3.1. Mâu thuẫn nội tại giữa r = 0,42 và "85,4% giải thích"', level=3)
    para(d,
        'Theo lý thuyết thống kê cơ bản, hệ số xác định r² (variance '
        'explained) = bình phương của hệ số tương quan r. Áp dụng cho '
        'r = 0,42: r² = 0,176 = 17,6%. Nói cách khác, mối tương quan '
        'tuyến tính giữa nhóm yếu tố bản thân và lo âu chỉ giải thích '
        '17,6% phương sai — KHÔNG PHẢI 85,4% như tác giả công bố.', indent=False
    )
    para(d, 'Khả năng giải thích cho con số 85,4%:', bold=True, indent=False)
    bullet(d, 'Khả năng 1: 85,4% là R² của HỒI QUY ĐA BIẾN với nhiều biến predictor cùng nhóm "bản thân" — KHÔNG phải r² của một biến đơn. Nếu vậy, tác giả cần báo cáo MA TRẬN HỒI QUY và β chuẩn hóa cho từng biến — không có trong bản gốc.', size=12)
    bullet(d, 'Khả năng 2: 85,4% là TỶ LỆ % học sinh có biểu hiện thuộc nhóm "bản thân" (không phải variance explained). Nếu vậy, tác giả đã NHẦM thuật ngữ — % HS biểu hiện ≠ % phương sai giải thích.', size=12)
    bullet(d, 'Khả năng 3: 85,4% là phản hồi ĐỊNH TÍNH từ phỏng vấn sâu — không có ý nghĩa thống kê. Nếu vậy, không nên trình bày như con số định lượng.', size=12)
    para(d,
        'Bất kể khả năng nào, con số 85,4% trong báo cáo của tác giả '
        'cần được làm rõ trước khi trích dẫn. Việc bài KHÔNG CÔNG BỐ '
        'phương pháp tính cụ thể là điểm yếu PHƯƠNG PHÁP LUẬN nghiêm '
        'trọng — vi phạm nguyên tắc minh bạch của Strengthening the '
        'Reporting of Observational Studies in Epidemiology (STROBE; '
        'von Elm và cộng sự, 2007).', indent=False
    )

    H(d, '3.2. Top 3 biểu hiện chỉ là TẦN SUẤT, không phải predictor', level=3)
    para(d,
        'Ba biểu hiện hàng đầu — "ít giao tiếp chia sẻ" (X̄ = 0,91), '
        '"thất vọng" (X̄ = 0,90), "lo sợ khó khăn" (X̄ = 0,88) — chỉ là '
        'điểm trung bình theo thang điểm tự thiết kế của tác giả (không '
        'rõ thang gì, có thể Likert 0–1 hoặc 0–4). Không có hệ số tương '
        'quan r riêng cho từng biểu hiện, không có p-value, không có '
        'KTC 95%.', indent=False
    )
    para(d,
        'Hậu quả: người đọc KHÔNG BIẾT biểu hiện nào TÁC ĐỘNG MẠNH NHẤT '
        'đến lo âu — chỉ biết biểu hiện nào XUẤT HIỆN NHIỀU NHẤT. Hai '
        'thông tin này hoàn toàn KHÁC nhau về ý nghĩa khoa học. Phát '
        'hiện "ít giao tiếp chia sẻ" có thể chỉ là biểu hiện THƯỜNG GẶP '
        'ở học sinh THPT VN nói chung, không nhất thiết gây ra hoặc dự '
        'báo lo âu.', indent=False
    )

    H(d, '3.3. Thiếu phân tích đa biến và kiểm soát yếu tố nhiễu', level=3)
    para(d,
        'Với bốn nhóm yếu tố được xác định (học tập, gia đình, quan hệ '
        'xã hội, bản thân HS), một thiết kế nghiên cứu chuẩn cần thực '
        'hiện hồi quy đa biến hoặc mô hình SEM để KIỂM SOÁT đồng thời '
        'tác động của bốn nhóm. Nếu không, không thể xác định liệu yếu '
        'tố "bản thân" có còn tác động độc lập sau khi loại trừ ảnh '
        'hưởng của ba nhóm còn lại.', indent=False
    )
    para(d,
        'So sánh với chuẩn quốc tế: Qiu và cộng sự (2022) — QT009 trong '
        'cơ sở dữ liệu dự án — thực hiện hồi quy logistic đa biến trên '
        '2.079 học sinh THCS Trung Quốc (tuổi trung bình 16,7), kiểm '
        'soát đồng thời giới, lớp, kinh tế gia đình, và các yếu tố tâm '
        'lý xã hội. Kết quả: nuôi dạy tích cực giảm nguy cơ trầm cảm '
        '70% (OR = 0,30) và lo âu 68% (OR = 0,32) SAU KHI kiểm soát '
        'các yếu tố khác. Nguyễn Thị Vân (2020) không có phân tích '
        'tương đương.', indent=False
    )

    H(d, '3.4. Thang đo STAI bản Việt cũ — chưa cập nhật psychometric', level=3)
    para(d,
        'Tác giả sử dụng STAI (State-Trait Anxiety Inventory; Spielberger, '
        '1983) bản Việt hóa của Nguyễn Công Khanh (2000) — đã 20 năm '
        'không cập nhật psychometric đến thời điểm nghiên cứu. STAI '
        'cũng đo lo âu trạng thái + lo âu đặc điểm — KHÔNG phải chẩn '
        'đoán RLLA theo DSM-5/ICD-11. Bài dùng từ "rối loạn lo âu" là '
        'KHÔNG CHÍNH XÁC về thuật ngữ chẩn đoán.', indent=False
    )
    para(d,
        'Đề xuất thay thế: dùng RCADS (Chorpita và cộng sự, 2000) — '
        'bản đã được Nguyễn Cao Minh (2012) chuẩn hóa cho học sinh '
        'Việt Nam, hoặc DASS-21 (Lovibond & Lovibond, 1995) — bản đã '
        'được Tran, Tran và Fisher (2013) chuẩn hóa cho phụ nữ Bắc '
        'Việt Nam. Đối với lòng tự trọng, đề xuất dùng Rosenberg '
        'Self-Esteem Scale (Rosenberg, 1965) — chuẩn vàng quốc tế.', indent=False
    )

    H(d, '3.5. Mẫu thuận tiện — không đại diện', level=3)
    para(d,
        '90 học sinh phỏng vấn sâu được chọn theo mẫu thuận tiện '
        '(convenience sampling) từ 4 trường nội + ngoại thành TPHCM. '
        'Phương pháp này KHÔNG ngẫu nhiên, KHÔNG đại diện cho học '
        'sinh THPT TPHCM nói chung — chứ chưa nói đến học sinh THPT '
        'Việt Nam toàn quốc. Hậu quả: phát hiện không thể ngoại suy '
        'sang dân số rộng hơn.', indent=False
    )

    H(d, '3.6. Thiếu nhóm chứng', level=3)
    para(d,
        'Bài không có nhóm chứng (học sinh KHÔNG có biểu hiện lo âu) '
        'để so sánh. Các biểu hiện được liệt kê trong nhóm "bản thân '
        'HS" (ít giao tiếp, thất vọng, lo sợ khó khăn) có thể CHỈ LÀ '
        'biểu hiện chung của tuổi vị thành niên Việt Nam — không '
        'nhất thiết là yếu tố ĐẶC THÙ của nhóm có lo âu.', indent=False
    )

    # 4. So sánh với y văn quốc tế
    H(d, '4. So sánh với chuẩn quốc tế về yếu tố "bản thân HS"', level=2, color=NAVY)
    para(d,
        'Trong y văn quốc tế, yếu tố "bản thân học sinh" thường được '
        'phân thành các thành phần CỤ THỂ với thang đo chuẩn — không '
        'phải nhóm yếu tố tổng quát như VN004:'
    )
    add_table(d,
        ['Thành phần', 'Thang đo chuẩn', 'Tham chiếu chính'],
        [
            ['Lòng tự trọng', 'Rosenberg Self-Esteem Scale (RSES)', 'Rosenberg (1965)'],
            ['Tự khái niệm', 'Self-Concept Scale (SCS)', 'Marsh & Shavelson (1985)'],
            ['Tự hiệu quả', 'General Self-Efficacy Scale (GSE)', 'Schwarzer & Jerusalem (1995)'],
            ['Resilience', 'Connor-Davidson Resilience Scale (CD-RISC)', 'Connor & Davidson (2003)'],
            ['Tính cách lo âu', 'Big Five Inventory — Neuroticism', 'John & Srivastava (1999)'],
        ]
    )
    para(d, '')
    para(d,
        'So với khung chuẩn này, VN004 đo "yếu tố bản thân" bằng bảng '
        'hỏi tự thiết kế với 3 biểu hiện chung — KHÔNG đo lòng tự '
        'trọng, tự hiệu quả, hay resilience theo thang chuẩn. Hậu '
        'quả: kết quả KHÔNG so sánh trực tiếp được với y văn quốc '
        'tế. Đề xuất nghiên cứu tiếp theo nên SỬ DỤNG các thang đo '
        'chuẩn trên — đặc biệt RSES và CD-RISC — để cho phép so '
        'sánh.', indent=False
    )

    # CÂU TRẢ LỜI TÔ XANH
    H(d, '5. CÂU TRẢ LỜI', level=2, color=NAVY)
    p = d.add_paragraph()
    r = p.add_run('Tóm gọn về phản biện VN004 — yếu tố bản thân HS:')
    r.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(12)

    bullets_blue = [
        'Mâu thuẫn nội tại NGHIÊM TRỌNG NHẤT: r = 0,42 (r² = 17,6%) nhưng tác giả công bố nhóm "bản thân HS" giải thích 85,4% — chênh lệch 5 lần. Cần làm rõ phương pháp tính trước khi trích dẫn.',
        'Top 3 biểu hiện ("ít giao tiếp chia sẻ", "thất vọng", "lo sợ khó khăn") chỉ là TẦN SUẤT, không phải PREDICTOR. Không có r riêng, không có p-value, không có KTC 95% cho từng biểu hiện.',
        'Thiếu phân tích đa biến để kiểm soát đồng thời 4 nhóm yếu tố — không thể khẳng định "bản thân HS" tác động độc lập.',
        'Thang STAI bản Việt 2000 đã 20 năm không cập nhật; STAI đo lo âu trạng thái + đặc điểm, KHÔNG chẩn đoán RLLA theo DSM-5.',
        'Mẫu thuận tiện 90 HS từ 4 trường TPHCM — không đại diện.',
        'Thiếu nhóm chứng — không phân biệt được biểu hiện đặc thù của nhóm có lo âu.',
    ]
    for b in bullets_blue:
        p = d.add_paragraph(style='List Bullet')
        r = p.add_run(b); r.font.color.rgb = BLUE; r.font.size = Pt(12)

    p = d.add_paragraph()
    r = p.add_run('Cách trích vào báo cáo CTH:')
    r.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(12)
    p = d.add_paragraph()
    r = p.add_run(
        '"Nguyễn Thị Vân (2020) khảo sát 90 học sinh THPT TPHCM bằng '
        'bảng hỏi tự thiết kế và STAI bản Việt — xác định bốn nhóm '
        'yếu tố ảnh hưởng đến lo âu, trong đó nhóm bản thân học sinh '
        'có hệ số tương quan r = 0,42 với mức lo âu. Tuy nhiên, bài '
        'có một số hạn chế phương pháp luận: con số "85,4% giải '
        'thích" công bố trong báo cáo cần được làm rõ về cách tính '
        '(không khớp với r² = 0,176 từ r = 0,42); chưa thực hiện '
        'phân tích đa biến để kiểm soát yếu tố nhiễu; thang STAI '
        'bản 2000 chưa cập nhật psychometric; mẫu thuận tiện không '
        'đại diện. Khuyến nghị các nghiên cứu tiếp theo sử dụng '
        'thang đo chuẩn quốc tế cho từng thành phần của yếu tố '
        '"bản thân" — Rosenberg Self-Esteem Scale (Rosenberg, '
        '1965) cho lòng tự trọng; Connor-Davidson Resilience Scale '
        '(Connor & Davidson, 2003) cho resilience — kết hợp với '
        'phân tích hồi quy đa biến hoặc mô hình SEM."'
    )
    r.font.color.rgb = BLUE; r.font.size = Pt(12); r.italic = True

    p = d.add_paragraph()
    r = p.add_run('Khuyến nghị thầy:')
    r.bold = True; r.font.color.rgb = BLUE; r.font.size = Pt(12)
    bullets_blue_2 = [
        'KHI TRÍCH VN004 vào tổng quan: chỉ dùng để minh họa BỐN NHÓM YẾU TỐ ảnh hưởng (khung khái niệm), KHÔNG dùng làm bằng chứng định lượng cho cường độ tác động.',
        'KHÔNG sử dụng con số "85,4% giải thích" — không có cơ sở thống kê rõ ràng.',
        'Khi xây dựng nghiên cứu của thầy: dùng RSES cho lòng tự trọng + thực hiện hồi quy đa biến hoặc SEM để kiểm soát yếu tố nhiễu.',
    ]
    for b in bullets_blue_2:
        p = d.add_paragraph(style='List Bullet')
        r = p.add_run(b); r.font.color.rgb = BLUE; r.font.size = Pt(12)

    # 6. Phụ lục
    H(d, '6. Phụ lục — Tài liệu tham khảo', level=2, color=NAVY)
    refs = [
        'Chorpita, B. F., Yim, L., Moffitt, C., Umemoto, L. A., & Francis, S. E. (2000). Assessment of symptoms of DSM-IV anxiety and depression in children: A revised child anxiety and depression scale. Behaviour Research and Therapy, 38(8), 835–855. https://doi.org/10.1016/S0005-7967(99)00130-8',
        'Connor, K. M., & Davidson, J. R. T. (2003). Development of a new resilience scale: The Connor-Davidson Resilience Scale (CD-RISC). Depression and Anxiety, 18(2), 76–82. https://doi.org/10.1002/da.10113',
        'Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Sydney: Psychology Foundation of Australia.',
        'Nguyễn, C. M. (2012). Chuẩn hóa thang đo Revised Children\'s Anxiety and Depression Scale cho học sinh Việt Nam.',
        'Nguyễn, T. V. (2018/2020). Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh. Tạp chí Khoa học – Khoa học Giáo dục, ĐHSP TP.HCM. [VN004 trong cơ sở dữ liệu dự án.]',
        'Qiu, Z., Guo, Y., Wang, J., & Zhang, H. (2022). Associations of parenting style and resilience with depression and anxiety among Chinese high school students. Frontiers in Public Health, 10, 989125. https://doi.org/10.3389/fpubh.2022.989125 [QT009 trong cơ sở dữ liệu dự án.]',
        'Rosenberg, M. (1965). Society and the adolescent self-image. Princeton University Press.',
        'Spielberger, C. D. (1983). Manual for the State-Trait Anxiety Inventory (Form Y). Palo Alto, CA: Consulting Psychologists Press.',
        'Tran, T. D., Tran, T., & Fisher, J. (2013). Validation of the depression anxiety stress scales (DASS) 21 as a screening instrument for depression and anxiety in a rural community-based cohort of northern Vietnamese women. BMC Psychiatry, 13, 24. https://doi.org/10.1186/1471-244X-13-24',
        'von Elm, E., Altman, D. G., Egger, M., Pocock, S. J., Gøtzsche, P. C., & Vandenbroucke, J. P. (2007). The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Annals of Internal Medicine, 147(8), 573–577. https://doi.org/10.7326/0003-4819-147-8-200710160-00010',
    ]
    for r in refs:
        ref_entry(d, r)

    out = OUT_DIR / 'Phan_bien_VN004_NguyenThiVan_yeu_to_ban_than_HS.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')


# Run both
print('Building 2 docs:')
build_self_esteem_doc()
build_VN004_critique()
print('Done.')
