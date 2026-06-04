# -*- coding: utf-8 -*-
"""
QA TOAN BO bao cao trong 01_Bao-cao/ — apply category A automated checks.

Output: QA matrix + danh sach issue ra json
"""
import os, sys, io, re, json
from collections import defaultdict, Counter
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
BAOCAO_DIR = os.path.join(ROOT, '01_Bao-cao')

def read_doc(path):
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    text = '\n'.join(paras)
    table_text = []
    n_tables = len(doc.tables)
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                if cell.text.strip():
                    table_text.append(cell.text)
    full = text + '\n' + '\n'.join(table_text)
    return doc, paras, full, n_tables

def check_internal_consistency(text):
    fact_pat = re.compile(
        r'(\w+(?:\s+(?:&|et al\.))?\s+\d{4})[^.]*?'
        r'(d\s*=\s*-?\d+[,\.]\d+|g\s*=\s*-?\d+[,\.]\d+|β\s*=\s*-?\d+[,\.]\d+|'
        r'OR\s*=\s*-?\d+[,\.]\d+|SMD\s*=\s*-?\d+[,\.]\d+|n\s*=\s*\d+[\.,]?\d*|'
        r'p\s*[<>]\s*0[,\.]\d+|SUCRA\s*=?\s*\d+[,\.]?\d*\s*%?)',
        re.IGNORECASE
    )
    facts = defaultdict(list)
    for m in fact_pat.finditer(text):
        facts[m.group(1).strip()].append(m.group(2).strip())
    issues = []
    for ay, fs in facts.items():
        by_stat = defaultdict(set)
        for f in fs:
            stype = re.match(r'(\w+|β)', f).group(1).lower()
            val_m = re.search(r'-?\d+[,\.]?\d*', f)
            if val_m:
                by_stat[stype].add(val_m.group(0).replace(',', '.'))
        for stype, vals in by_stat.items():
            if len(vals) > 1:
                issues.append(f'{ay} {stype} = {vals}')
    return issues

def check_number_plausibility(text):
    issues = []
    # %
    for m in re.finditer(r'(\d+[,\.]\d+)\s*%', text):
        try:
            val = float(m.group(1).replace(',', '.'))
            if val > 100:
                issues.append(f'PCT > 100: {m.group(0)}')
        except: pass
    # p-values
    for m in re.finditer(r'p\s*[<>=]\s*(\d+[,\.]?\d*)', text, re.IGNORECASE):
        try:
            val = float(m.group(1).replace(',', '.'))
            if val > 1:
                issues.append(f'p > 1: {m.group(0)}')
        except: pass
    # d/g extreme
    for m in re.finditer(r'(?:Cohen[\'’]?s?\s+d|Hedges[\'’]?s?\s+g)\s*=\s*(-?\d+[,\.]\d+)', text, re.IGNORECASE):
        try:
            val = abs(float(m.group(1).replace(',', '.')))
            if val > 3:
                issues.append(f'd/g extreme: {m.group(0)}')
        except: pass
    return issues

def check_citation_completeness(paras):
    """Find paragraphs with statistics but no inline citation"""
    issues = []
    stat_kw = ['Cohen d', 'Hedges g', 'OR =', 'SMD =', 'SUCRA', 'AOR =', 'NNT', 'β =']
    author_pat = re.compile(r'(\w[\w-]{2,}(?:\s+(?:&|et al\.|và cộng sự))?\s*\(?\d{4}\)?)')
    for i, para in enumerate(paras):
        has_stat = any(kw in para for kw in stat_kw)
        if not has_stat and re.search(r'\d+[,\.]\d+\s*%', para):
            has_stat = True
        if has_stat and len(para) > 60:
            has_cite = bool(author_pat.search(para)) or 'CAMS' in para or 'NEJM' in para
            if not has_cite:
                issues.append(f'#{i}: {para[:100]}')
    return issues

def check_table_format(doc):
    issues = []
    for ti, tb in enumerate(doc.tables, start=1):
        if len(tb.rows) < 2:
            continue
        cols = [len(r.cells) for r in tb.rows]
        if len(set(cols)) > 1:
            issues.append(f'Bảng {ti}: cột không đều {cols}')
        empty_headers = sum(1 for c in tb.rows[0].cells if not c.text.strip())
        if empty_headers > 0:
            issues.append(f'Bảng {ti}: {empty_headers} header trống')
        empty_cells = sum(1 for row in tb.rows for c in row.cells if not c.text.strip())
        total = sum(len(r.cells) for r in tb.rows)
        if total > 0 and empty_cells / total > 0.4:
            issues.append(f'Bảng {ti}: {empty_cells}/{total} cells trống ({empty_cells*100//total}%)')
    return issues

def check_year_consistency(text):
    author_years = defaultdict(set)
    ay_pat = re.compile(r'\b(\w[\w-]+)\s+(?:&\s+\w+\s+)?(?:et al\.\s+)?(\d{4})\b')
    skip = {'bài', 'phần', 'mục', 'bảng', 'trang', 'năm', 'từ', 'sau', 'trước',
            'page', 'pages', 'vol', 'volume', 'số'}
    for m in ay_pat.finditer(text):
        author = m.group(1)
        year = m.group(2)
        if author.lower() in skip or len(author) < 4 or not author[0].isupper():
            continue
        author_years[author].add(year)
    issues = []
    for a, ys in author_years.items():
        if len(ys) > 1:
            issues.append(f'{a}: {sorted(ys)}')
    return issues

# ============================================================
# RUN ON ALL REPORTS
# ============================================================
print('=' * 80)
print('QA TOÀN BỘ — Multi-layer automated check on all reports')
print('=' * 80)

reports = sorted([f for f in os.listdir(BAOCAO_DIR) if f.endswith('.docx')])
print(f'\nTổng số báo cáo: {len(reports)}')

results = {}
for r in reports:
    path = os.path.join(BAOCAO_DIR, r)
    try:
        doc, paras, full_text, n_tables = read_doc(path)
    except Exception as e:
        print(f'\n[ERR] {r}: {e}')
        continue

    print(f'\n--- {r[:60]} ---')
    print(f'    {len(paras)} đoạn, {n_tables} bảng, {len(full_text):,} ký tự')

    c1 = check_internal_consistency(full_text)
    c2 = check_number_plausibility(full_text)
    c3 = check_citation_completeness(paras)
    c4 = check_table_format(doc)
    c5 = check_year_consistency(full_text)

    total = len(c1) + len(c2) + len(c3) + len(c4) + len(c5)
    status = '✓ OK' if total == 0 else f'⚠ {total} issues'
    print(f'    [{status}] consistency={len(c1)} plausibility={len(c2)} '
          f'cite={len(c3)} table={len(c4)} year={len(c5)}')

    results[r] = {
        'paragraphs': len(paras),
        'tables': n_tables,
        'chars': len(full_text),
        'check1_consistency': c1,
        'check2_plausibility': c2,
        'check3_citation': c3,
        'check4_table': c4,
        'check5_year': c5,
        'total_issues': total,
    }

# ============================================================
# Summary matrix
# ============================================================
print('\n' + '=' * 80)
print('MA TRẬN QA — Tổng hợp')
print('=' * 80)

print(f'\n{"Báo cáo":<55} {"Đoạn":>6} {"Bảng":>5} {"Issues":>7}')
print('-' * 80)
for r, info in results.items():
    name = r[:55]
    print(f'{name:<55} {info["paragraphs"]:>6} {info["tables"]:>5} {info["total_issues"]:>7}')

total_issues = sum(info['total_issues'] for info in results.values())
total_reports_clean = sum(1 for info in results.values() if info['total_issues'] == 0)
print('-' * 80)
print(f'Tổng issues: {total_issues}')
print(f'Báo cáo sạch: {total_reports_clean}/{len(results)}')

# Save detailed log
log_path = os.path.join(os.path.dirname(__file__), 'qa_all_reports_log.json')
with open(log_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, ensure_ascii=False, indent=2)
print(f'\nChi tiết JSON: {log_path}')

# Show top issues
print('\n=== TOP REAL ISSUES (cần kiểm tra) ===')
real_issues = []
for r, info in results.items():
    for issue_type in ['check1_consistency', 'check5_year']:
        for i in info[issue_type]:
            real_issues.append((r, issue_type, i))

if real_issues:
    print(f'Total: {len(real_issues)} real consistency/year issues')
    for r, t, i in real_issues[:15]:
        print(f'  [{r[:30]}] [{t.replace("check", "")}] {i[:120]}')
else:
    print('Không có real critical issues!')
