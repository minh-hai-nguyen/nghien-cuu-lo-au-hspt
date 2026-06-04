"""Build full bilingual 1-1 translation for QT067 Pascoe et al. 2020.
Each English paragraph is followed by Vietnamese translation.
References list is bilingual (English original + Vietnamese gloss).
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/03_Ban-dich/QT067_Pascoe_AcademicStress_IJAY_2020.docx')

doc = Document()

for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x30)
GREY = RGBColor(0x66, 0x66, 0x66)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def page_marker(n):
    p = doc.add_paragraph()
    r = p.add_run(f'──────── Trang {n} ────────')
    r.bold = True; r.font.size = Pt(11); r.font.color.rgb = GREY
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return p

def en(text):
    p = doc.add_paragraph()
    r = p.add_run('[EN] ')
    r.bold = True; r.font.color.rgb = GREY; r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(11); r2.italic = True
    return p

def vn(text):
    p = doc.add_paragraph()
    r = p.add_run('[VI] ')
    r.bold = True; r.font.color.rgb = NAVY; r.font.size = Pt(10)
    r2 = p.add_run(text)
    r2.font.size = Pt(12)
    return p

def pair(en_text, vn_text):
    en(en_text); vn(vn_text)

def note(text, color=RED):
    p = doc.add_paragraph()
    r = p.add_run('★ ' + text); r.italic = True
    r.font.color.rgb = color; r.font.size = Pt(11)
    return p

def ref_pair(en_text, vn_gloss):
    p1 = doc.add_paragraph()
    r1 = p1.add_run(en_text); r1.font.size = Pt(10)
    p2 = doc.add_paragraph()
    r2 = p2.add_run('  → ' + vn_gloss); r2.italic = True
    r2.font.color.rgb = NAVY; r2.font.size = Pt(10)

# ============================================================
# COVER + METADATA
# ============================================================
H('QT067 — Pascoe, Hetrick & Parker (2020)', level=1)
H('The impact of stress on students in secondary school and higher education', level=2)
H('Tác động của stress đối với học sinh trung học và sinh viên đại học', level=2, color=NAVY)

p = doc.add_paragraph()
r = p.add_run('International Journal of Adolescence and Youth, 25(1), 104–112.\nDOI: 10.1080/02673843.2019.1596823')
r.italic = True; r.font.size = Pt(11)

H('Thông tin định danh', level=2)
meta = [
    ('Mã canonical', 'QT067'),
    ('Tác giả', 'Michaela C. Pascoe (a,b), Sarah E. Hetrick (c,d), Alexandra G. Parker (a,c)'),
    ('Cơ quan', '(a) Institute for Health and Sport, Victoria University, Melbourne, Australia\n(b) Department of Cancer Experiences, Peter MacCallum Cancer Centre, Melbourne, Australia\n(c) Orygen, the National Centre of Excellence in Youth Mental Health, University of Melbourne, Australia\n(d) Department of Psychological Medicine, University of Auckland, New Zealand'),
    ('Năm', '2020 (Received 29 Jan 2019, Accepted 14 Mar 2019, Published online 11 Apr 2019)'),
    ('Tạp chí', 'International Journal of Adolescence and Youth (IJAY)'),
    ('Loại bài', 'Narrative review (tổng quan tường thuật)'),
    ('Đối tượng', 'Học sinh trung học (secondary) + sinh viên đại học (tertiary)'),
    ('License', 'Open Access (CC BY 4.0)'),
    ('Email tác giả', 'Michaela.pascoe@vu.edu.au'),
    ('Bài gốc', '02_Papers-goc/The-gioi_Au-My-Uc/QT067_Pascoe_AcademicStress_IJAY_2020.pdf'),
]
tbl = doc.add_table(rows=len(meta), cols=2)
tbl.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(meta):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    for cell_idx, cell in enumerate(tbl.rows[i].cells):
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(10); r.font.name = 'Times New Roman'
                if cell_idx == 0:
                    r.bold = True

note('Lưu ý: Bài Pascoe 2020 KHÔNG có bảng số liệu, KHÔNG có hình ảnh trong bản gốc. Toàn bộ là văn bản narrative review + ~80 references. Bản dịch dưới đây dịch song song 1-1 paragraph-by-paragraph theo Nguyên tắc dịch v2.', color=GREEN)

# ============================================================
# TRANG 1 — COVER VU REPOSITORY
# ============================================================
page_marker(1)
H('Bìa lưu trữ Victoria University Repository', level=2)
pair(
    "The impact of stress on students in secondary school and higher education. This is the Published version of the following publication: Pascoe, Michaela, Hetrick, Sarah and Parker, Alexandra (2019) The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth. ISSN 0267-3843. The publisher's official version can be found at https://www.tandfonline.com/doi/full/10.1080/02673843.2019.1596823. Note that access to this version may require subscription. Downloaded from VU Research Repository https://vuir.vu.edu.au/39399/",
    "Tác động của stress đối với học sinh trung học và sinh viên đại học. Đây là phiên bản đã xuất bản của công trình sau: Pascoe, Michaela, Hetrick, Sarah và Parker, Alexandra (2019) Tác động của stress đối với học sinh trung học và sinh viên đại học. Tạp chí Quốc tế về Vị thành niên và Thanh niên. ISSN 0267-3843. Phiên bản chính thức của nhà xuất bản có thể tìm thấy tại https://www.tandfonline.com/doi/full/10.1080/02673843.2019.1596823. Lưu ý rằng truy cập phiên bản đó có thể cần đăng ký. Tải xuống từ Kho lưu trữ Nghiên cứu Đại học Victoria https://vuir.vu.edu.au/39399/."
)

# ============================================================
# TRANG 2 — JOURNAL HOMEPAGE INFO
# ============================================================
page_marker(2)
H('Thông tin tạp chí (T&F)', level=2)
pair(
    "International Journal of Adolescence and Youth. ISSN: 0267-3843 (Print) 2164-4527 (Online). Journal homepage: https://www.tandfonline.com/loi/rady20. The impact of stress on students in secondary school and higher education. Michaela C. Pascoe, Sarah E. Hetrick & Alexandra G. Parker. To cite this article: Michaela C. Pascoe, Sarah E. Hetrick & Alexandra G. Parker (2019): The impact of stress on students in secondary school and higher education, International Journal of Adolescence and Youth, DOI: 10.1080/02673843.2019.1596823. © 2019 The Author(s). Published by Informa UK Limited, trading as Taylor & Francis Group. Published online: 11 Apr 2019. Article views: 7885.",
    "Tạp chí Quốc tế về Vị thành niên và Thanh niên. ISSN: 0267-3843 (bản in) 2164-4527 (bản trực tuyến). Trang chủ tạp chí: https://www.tandfonline.com/loi/rady20. Tác động của stress đối với học sinh trung học và sinh viên đại học. Michaela C. Pascoe, Sarah E. Hetrick & Alexandra G. Parker. Để trích dẫn bài báo này: Michaela C. Pascoe, Sarah E. Hetrick & Alexandra G. Parker (2019): Tác động của stress đối với học sinh trung học và sinh viên đại học, Tạp chí Quốc tế về Vị thành niên và Thanh niên, DOI: 10.1080/02673843.2019.1596823. © 2019 (các) Tác giả. Xuất bản bởi Informa UK Limited, hoạt động dưới tên Taylor & Francis Group. Trực tuyến từ ngày 11/04/2019. Lượt xem: 7885."
)

# ============================================================
# TRANG 3 — TITLE + AUTHORS + ABSTRACT + SECTION 1
# ============================================================
page_marker(3)

H('Tác giả & cơ quan', level=2)
pair(
    "Michaela C. Pascoe (a,b), Sarah E. Hetrick (c,d) and Alexandra G. Parker (a,c). (a) Institute for Health and Sport, Victoria University, Melbourne, Australia; (b) Department of Cancer Experiences, Peter MacCallum Cancer Centre, Melbourne, Australia; (c) Orygen, the National Centre of Excellence in Youth Mental Health (Centre for Youth Mental Health, University of Melbourne) University of Melbourne, Australia; (d) Department of Psychological Medicine, Faculty of Medical and Health Sciences, University of Auckland, Auckland, New Zealand.",
    "Michaela C. Pascoe (a,b), Sarah E. Hetrick (c,d) và Alexandra G. Parker (a,c). (a) Viện Sức khỏe và Thể thao, Đại học Victoria, Melbourne, Úc; (b) Khoa Trải nghiệm Ung thư, Trung tâm Ung thư Peter MacCallum, Melbourne, Úc; (c) Orygen, Trung tâm Quốc gia Xuất sắc về Sức khỏe Tâm thần Thanh niên (Trung tâm Sức khỏe Tâm thần Thanh niên, Đại học Melbourne), Đại học Melbourne, Úc; (d) Khoa Y học Tâm thần, Khoa Y và Khoa học Sức khỏe, Đại học Auckland, Auckland, New Zealand."
)

H('ABSTRACT — TÓM TẮT', level=2)
pair(
    "Students in secondary and tertiary education settings face a wide range of ongoing stressors related to academic demands. Previous research indicates that academic-related stress can reduce academic achievement, decrease motivation and increase the risk of school dropout. The longer-term impacts, which include reduced likelihood of sustainable employment, cost Governments billions of dollars each year. This narrative review presents the most recent research concerning the impact of academic-related stress, including discussion of the impact on students' learning capacity and academic performance, mental health problems, such as depression and anxiety, sleep disturbances and substance use.",
    "Học sinh ở các bậc trung học và đại học đối mặt với nhiều áp lực liên tục liên quan đến yêu cầu học tập. Các nghiên cứu trước đây cho thấy stress liên quan đến học tập có thể làm giảm thành tích học tập, giảm động lực và tăng nguy cơ bỏ học. Các tác động dài hạn — bao gồm khả năng việc làm bền vững giảm — gây thiệt hại hàng tỷ đô-la mỗi năm cho các Chính phủ. Bài tổng quan tường thuật này trình bày các nghiên cứu mới nhất về tác động của stress học đường, bao gồm thảo luận về ảnh hưởng đến năng lực học tập và thành tích học tập của học sinh, các vấn đề sức khỏe tâm thần như trầm cảm và lo âu, rối loạn giấc ngủ và sử dụng chất gây nghiện."
)

H('Article history & Keywords', level=3)
pair(
    "ARTICLE HISTORY: Received 29 January 2019 / Accepted 14 March 2019. KEYWORDS: Academic; adolescent health; education; mental health and well-being; stress.",
    "LỊCH SỬ BÀI BÁO: Nhận 29/01/2019 / Chấp nhận 14/03/2019. TỪ KHÓA: Học tập; sức khỏe vị thành niên; giáo dục; sức khỏe tâm thần và an sinh; stress."
)

H('Young people report high levels of stress — Người trẻ báo cáo mức stress cao', level=2)
pair(
    "Students in secondary and tertiary education settings face a wide range of ongoing normative stressors, which can be defined as normal day to day hassles such as ongoing academic demands. Accordingly, secondary/high school (defined here as junior/lower secondary education and senior/upper secondary education) (UNESCO, 2012) and tertiary (defined here as post-secondary education) (UNESCO, 2012) students commonly self-report experiencing ongoing stress relating to their education, which we refer to as academic-related stress, such as pressure to achieve high marks and concerns about receiving poor grades.",
    "Học sinh trung học và sinh viên đại học đối mặt với nhiều stress thường xuyên có tính chất chuẩn mực (normative stressors) — định nghĩa là các phiền toái hàng ngày bình thường, chẳng hạn yêu cầu học tập liên tục. Vì vậy, học sinh trung học (ở đây bao gồm THCS và THPT theo UNESCO, 2012) và sinh viên đại học (post-secondary, UNESCO, 2012) thường tự báo cáo trải qua stress liên quan đến học tập (academic-related stress), bao gồm áp lực phải đạt điểm cao và lo lắng về điểm kém."
)

pair(
    "For example, the Organisation for Economic Co-operation and Development (OECD) recently conducted a survey involving 72 countries and consisting of 540,000 student respondents aged 15–16 years. On average across OECD countries, 66% of students reported feeling stressed about poor grades and 59% reported that they often worry that taking a test will be difficult. The OECD further found that 55% of students feel very anxious about school testing, even when they are well prepared. As many 37% of students reported feeling very tense when studying, with girls consistently reporting greater anxiety relating to schoolwork compared to boys (OECD, 2017).",
    "Ví dụ, Tổ chức Hợp tác và Phát triển Kinh tế (OECD) gần đây thực hiện khảo sát ở 72 quốc gia với 540.000 học sinh 15–16 tuổi. Trung bình các nước OECD, 66% học sinh báo cáo bị stress vì điểm kém và 59% thường lo lắng rằng bài kiểm tra sẽ khó. OECD còn thấy 55% học sinh rất lo lắng về kiểm tra ở trường, ngay cả khi đã chuẩn bị kỹ. Có tới 37% học sinh báo cáo cảm thấy rất căng thẳng khi học, với nữ liên tục báo cáo lo âu liên quan tới bài vở cao hơn nam (OECD, 2017)."
)

pair(
    "This data demonstrates that education and academic performance are a significant source of stress to students. The impact of this ongoing academic-related stress to student outcomes and well-being has not been comprehensibly explored. Therefore, the current narrative review explores the impact of academic-related stress on students' academic performance, mental health and well-being.",
    "Dữ liệu này chứng tỏ giáo dục và thành tích học tập là nguồn stress đáng kể đối với học sinh. Tác động của stress học đường liên tục lên kết quả và an sinh (well-being) của học sinh chưa được khám phá toàn diện. Do đó, bài tổng quan tường thuật này khảo sát tác động của stress học đường lên thành tích học tập, sức khỏe tâm thần và an sinh của học sinh."
)

# ============================================================
# TRANG 4 — METHODS + MENTAL HEALTH
# ============================================================
page_marker(4)
H('Methods — Phương pháp', level=2)
pair(
    "A single author (MP) searched PubMed and Google Scholar for peer-reviewed articles published at any time in English. Search terms included academic, school, university, stress, mental health, depression, anxiety, youth, young people, resilience, stress management, stress education, substance use, sleep, drop-out, physical health with a combination of any and/or all of the preceding terms. A snowball strategy allowed for examination of references in identified articles, and inclusion of additional articles as appropriate. The author reviewed all potential articles for inclusion. Articles from all countries were included in this narrative review, if a school based (secondary [as defined at grade 7 or higher] or university) population was included and the study assessed the impact of stress on student mental health, substance use, sleep, dropout rates, physical activity or academic outcomes. Articles were included regardless of study design.",
    "Một tác giả duy nhất (MP) tìm trên PubMed và Google Scholar các bài báo bình duyệt xuất bản tại bất kỳ thời điểm nào, bằng tiếng Anh. Từ khóa tìm kiếm bao gồm academic, school, university, stress, mental health, depression, anxiety, youth, young people, resilience, stress management, stress education, substance use, sleep, drop-out, physical health với bất kỳ kết hợp nào của các từ trên. Chiến lược 'quả cầu tuyết' (snowball) cho phép xem references trong bài đã tìm được, và đưa thêm các bài phù hợp. Tác giả xem xét toàn bộ các bài có khả năng phù hợp. Bài từ mọi quốc gia đều được đưa vào nếu sử dụng dân số dựa trên trường học (trung học [định nghĩa từ lớp 7 trở lên] hoặc đại học) và đánh giá tác động của stress lên sức khỏe tâm thần, sử dụng chất, giấc ngủ, tỷ lệ bỏ học, hoạt động thể chất hoặc kết quả học tập. Bài được đưa vào bất kể thiết kế nghiên cứu."
)

note('PHẢN BIỆN: Phương pháp KHÔNG ĐẠT chuẩn systematic review — chỉ 1 tác giả tìm, không PRISMA, không có inclusion/exclusion criteria rõ ràng, không đánh giá chất lượng (risk of bias). Đây là điểm yếu lớn — không thể coi đây là evidence summary tin cậy.')

H('Academic-related stress and mental health — Stress học đường và sức khỏe tâm thần', level=2)
pair(
    "Previous research indicates that self-reported stress is associated with the presentation of anxious states and lower well-being (Carter, Garber, Ciesla, & Cole, 2006; Kessler, 1997; Robotham & Julian, 2006). The recent above-mentioned OECD survey reports that secondary students who self-report higher levels of academic-related stress also report lower well-being, measured using psychological, social, cognitive and physical components (OECD, 2015). A systematic review of 13 studies showed that in individuals undertaking higher education, self-reported levels of stress are associated with poorer quality of life and well-being (Ribeiro et al., 2017).",
    "Các nghiên cứu trước đây cho thấy stress tự báo cáo có liên quan đến biểu hiện lo âu và an sinh thấp hơn (Carter, Garber, Ciesla, & Cole, 2006; Kessler, 1997; Robotham & Julian, 2006). Khảo sát OECD đã đề cập ở trên báo cáo rằng học sinh trung học tự báo cáo mức stress học đường cao hơn cũng báo cáo an sinh thấp hơn — đo bằng các thành phần tâm lý, xã hội, nhận thức và thể chất (OECD, 2015). Một tổng quan hệ thống 13 nghiên cứu cho thấy ở người đang học bậc cao đẳng/đại học, mức stress tự báo cáo có liên quan đến chất lượng cuộc sống và an sinh kém hơn (Ribeiro và cs., 2017)."
)

pair(
    "Ongoing stress also precipitates the development of more serious mental health issues such as anxiety and depression (Kessler, 1997; Moylan, Maes, Wray, & Berk, 2013). The prevalence of anxiety is as high as 35% in tertiary students (Bayram & Bilgel, 2008; Eisenberg, Gollust, Golberstein, & Hefner, 2007; Ozen, Ercan, Irgil, & Sigirli, 2010) and the prevalence of depression is 30% (Ibrahim, Kelly, Adams, & Glazebrook, 2013). The reciprocal relationship between stress and depression and anxiety is well established (Dantzer, 2012; Dantzer, O'Connor, Lawson, & Kelley, 2011; Maes, 2008). Indeed, major stressful life events are one of the best predictors of the onset of depression (Kendler et al., 1995; Kessler, 1997). Accordingly, in young people the first onset of depression is often preceded by major life stressors (Lewinsohn, Allen, Seeley, & Gotlib, 1999).",
    "Stress kéo dài cũng thúc đẩy sự phát triển các vấn đề sức khỏe tâm thần nghiêm trọng hơn như lo âu và trầm cảm (Kessler, 1997; Moylan, Maes, Wray, & Berk, 2013). Tỷ lệ lo âu lên tới 35% ở sinh viên đại học (Bayram & Bilgel, 2008; Eisenberg, Gollust, Golberstein, & Hefner, 2007; Ozen, Ercan, Irgil, & Sigirli, 2010) và tỷ lệ trầm cảm là 30% (Ibrahim, Kelly, Adams, & Glazebrook, 2013). Mối quan hệ hai chiều giữa stress và trầm cảm/lo âu đã được xác lập rõ (Dantzer, 2012; Dantzer, O'Connor, Lawson, & Kelley, 2011; Maes, 2008). Thực tế, các biến cố căng thẳng lớn trong đời là một trong những yếu tố dự báo tốt nhất cho khởi phát trầm cảm (Kendler và cs., 1995; Kessler, 1997). Do đó ở người trẻ, lần khởi phát trầm cảm đầu tiên thường có các stressor lớn đi trước (Lewinsohn, Allen, Seeley, & Gotlib, 1999)."
)

pair(
    "Aside from impairing overall health and well-being, depression and anxiety symptoms can further adversely affect academic achievement (Bernal-Morales, Rodríguez-Landa, & Pulido-Criollo, 2015). In undergraduate university students from the United States, those with higher self-reported anxiety and depression symptoms were found to achieve poorer grades on examinations (Chapell et al., 2005; Hysenbegasi, Hass, & Rowland, 2005). A longitudinal study of Hawaiian secondary school students showed that self-reported depressive symptoms resulted in subsequent poor academic achievement (Kessler, 2012; McArdle, Hamagami, Chang, & Hishinuma, 2014).",
    "Bên cạnh việc làm suy giảm sức khỏe và an sinh tổng thể, các triệu chứng trầm cảm và lo âu còn ảnh hưởng tiêu cực đến thành tích học tập (Bernal-Morales, Rodríguez-Landa, & Pulido-Criollo, 2015). Ở sinh viên đại học hệ cử nhân tại Mỹ, những người tự báo cáo triệu chứng lo âu và trầm cảm cao hơn có điểm kém hơn trong các kỳ thi (Chapell và cs., 2005; Hysenbegasi, Hass, & Rowland, 2005). Một nghiên cứu thuần tập (longitudinal) trên học sinh trung học ở Hawaii cho thấy các triệu chứng trầm cảm tự báo cáo dẫn đến thành tích học kém về sau (Kessler, 2012; McArdle, Hamagami, Chang, & Hishinuma, 2014)."
)

pair(
    "This is consistent with the findings of Humensky et al. (2010) who found that self-reported depressive symptoms were associated with concentration difficulties and trouble with completing school tasks, in 83 students from the United States between the ages of 14–21, and at-risk for major depression (Humensky et al., 2010). In a sample of Finnish students aged 13–17, self-reported depression severity was associated with concentration difficulties, and poorer social relationships, self-learning, poorer academic performance, and worse reading and writing outcomes (Fröjd et al., 2008). Therefore, it is not surprising that young people with depression, particularly males, are less likely to undertake higher education, as shown in a 15-year longitudinal study of Swedish adolescents (Jonsson et al., 2010). Importantly, adolescent depression can also result in longer-term poor employment outcomes, as demonstrated by a 25-year longitudinal study of New Zealand children (n = 982). This study found that people who had depression at ages 16–21 had greater rates of welfare dependence and unemployment, demonstrating that the impact of poor mental health in adolescence can have long-lasting impacts",
    "Điều này nhất quán với phát hiện của Humensky và cs. (2010) — họ thấy các triệu chứng trầm cảm tự báo cáo có liên quan đến khó tập trung và khó hoàn thành nhiệm vụ học tập, ở 83 học sinh Mỹ tuổi 14–21 có nguy cơ trầm cảm chính (major depression) (Humensky và cs., 2010). Trong mẫu học sinh Phần Lan tuổi 13–17, mức độ nghiêm trọng của trầm cảm tự báo cáo có liên quan đến khó tập trung, mối quan hệ xã hội kém hơn, tự học kém hơn, thành tích học tập kém hơn và kết quả đọc viết tệ hơn (Fröjd và cs., 2008). Do đó, không ngạc nhiên khi người trẻ bị trầm cảm — đặc biệt là nam — ít theo học bậc đại học hơn, như nghiên cứu thuần tập 15 năm trên thanh thiếu niên Thụy Điển cho thấy (Jonsson và cs., 2010). Quan trọng hơn, trầm cảm vị thành niên còn dẫn đến kết quả việc làm kém về lâu dài, như nghiên cứu thuần tập 25 năm trên trẻ em New Zealand (n = 982) chứng minh. Nghiên cứu này phát hiện những người bị trầm cảm ở tuổi 16–21 có tỷ lệ phụ thuộc trợ cấp xã hội và thất nghiệp cao hơn, chứng tỏ tác động của sức khỏe tâm thần kém ở vị thành niên có thể kéo dài rất lâu"
)

# ============================================================
# TRANG 5 — MENTAL HEALTH cont + SUBSTANCE USE + SLEEP
# ============================================================
page_marker(5)

pair(
    "(Fergusson, Boden, & Horwood, 2007). Enhancing support in the education setting may improve the mental health of young people. A national telephone survey of United States households showed that the incidence of depression in college students decreases if students have positive adjustments to academic life as well as adequate social support (Ross & Mirowsky, 2006). Indeed, an Australian randomised control trial reported that a gamified online cognitive behaviour therapy intervention was effective in reducing depressive symptoms in 540 final year secondary students (Perry et al., 2017). This study demonstrates the potential of education settings in mediating the impacts of academic-related stress on young people's mental health.",
    "(Fergusson, Boden, & Horwood, 2007). Tăng cường hỗ trợ trong môi trường giáo dục có thể cải thiện sức khỏe tâm thần của người trẻ. Một khảo sát điện thoại toàn quốc tại các hộ gia đình Mỹ cho thấy tỷ lệ trầm cảm ở sinh viên đại học giảm nếu sinh viên có điều chỉnh tích cực với đời sống học tập và đủ hỗ trợ xã hội (Ross & Mirowsky, 2006). Thực tế, một thử nghiệm ngẫu nhiên có đối chứng (RCT) ở Úc báo cáo rằng can thiệp CBT trực tuyến dạng gamification có hiệu quả giảm triệu chứng trầm cảm ở 540 học sinh năm cuối THPT (Perry và cs., 2017). Nghiên cứu này chứng tỏ tiềm năng của môi trường giáo dục trong việc làm trung gian ảnh hưởng của stress học đường lên sức khỏe tâm thần thanh thiếu niên."
)

H('Academic-related stress and substance use — Stress học đường và sử dụng chất gây nghiện', level=2)
pair(
    "The health and risk behaviours of young people, including substance use and abuse, are all important determinants of their current and future health and well-being status (Tountas & Dimitrakaki, 2006; World Health Organisation, 2004). Academic-related stress can increase substance use among young people. In a survey study of 128 Grade 11 students attending competitive private schools in the United States, students who reported experiencing high ongoing stress, particularly in relation to academic achievement and the tertiary education admissions process, also reported high rates of drug and alcohol use (Leonard et al., 2015). The authors report that substance use was associated with a greater desire for academic achievement, higher perceived stress, less effective coping strategies, and less closeness with parents (Leonard et al., 2015).",
    "Hành vi sức khỏe và hành vi nguy cơ ở người trẻ — bao gồm sử dụng và lạm dụng chất gây nghiện — đều là yếu tố quyết định quan trọng đến tình trạng sức khỏe và an sinh hiện tại và tương lai của họ (Tountas & Dimitrakaki, 2006; WHO, 2004). Stress học đường có thể làm tăng sử dụng chất ở người trẻ. Trong nghiên cứu khảo sát 128 học sinh lớp 11 tại các trường tư cạnh tranh ở Mỹ, học sinh báo cáo trải qua stress cao kéo dài — đặc biệt liên quan đến thành tích học tập và quy trình tuyển sinh đại học — cũng báo cáo tỷ lệ sử dụng ma túy và rượu cao (Leonard và cs., 2015). Các tác giả báo cáo việc sử dụng chất có liên quan đến mong muốn thành tích cao hơn, stress nhận thức cao hơn, chiến lược ứng phó kém hiệu quả hơn và ít gần gũi với cha mẹ hơn (Leonard và cs., 2015)."
)

pair(
    "In 7th and 8th Grade students from the United States, self-perceived stress has similarly been reported to be related to substance use. In these students, coping strategies that included information gathering, problem solving and having a positive outlook, as well as adult social support and relaxation were inversely related to substance use (Wills, 1986). This study demonstrates the importance of protective social factors in mediating the effects of academic-related stress. In a cross-sectional study of tertiary nursing students from the United States, those with higher self-reported stress had higher incidence of substance use. Students who had higher perceptions of faculty support used fewer stimulants to assist them while studying, further demonstrating the proactive role of social factors (Boulton & O'Connell, 2017).",
    "Ở học sinh lớp 7-8 tại Mỹ, stress tự nhận thức cũng được báo cáo có liên quan đến sử dụng chất. Ở những học sinh này, các chiến lược ứng phó bao gồm thu thập thông tin, giải quyết vấn đề và có cái nhìn tích cực, cũng như hỗ trợ xã hội từ người lớn và thư giãn — có tương quan nghịch với sử dụng chất (Wills, 1986). Nghiên cứu này chứng tỏ tầm quan trọng của các yếu tố xã hội bảo vệ trong việc làm trung gian các tác động của stress học đường. Trong nghiên cứu cắt ngang trên sinh viên điều dưỡng đại học ở Mỹ, những người có stress tự báo cáo cao hơn có tỷ lệ sử dụng chất cao hơn. Sinh viên có nhận thức cao hơn về hỗ trợ từ giảng viên thì sử dụng ít chất kích thích để hỗ trợ học tập hơn, càng chứng tỏ vai trò chủ động của các yếu tố xã hội (Boulton & O'Connell, 2017)."
)

pair(
    "Finally, The Canadian Institute for Health reports that young people aged 12–19 who feel connected to their school report less anxiety and less risky behaviours, such as smoking and drinking alcohol, compared to those who do not feel connected to their schools (Canadian Institute for Health, 2005). Collectively, the above discussed findings indicate that increased stress is associated with substance use among students and that perceived social support, including from within the education environment, may positively mediate this relationship.",
    "Cuối cùng, Viện Sức khỏe Canada báo cáo rằng người trẻ tuổi 12-19 cảm thấy gắn kết với trường học của mình thì báo cáo lo âu thấp hơn và ít hành vi nguy cơ hơn — như hút thuốc và uống rượu — so với những người không cảm thấy gắn kết với trường (Canadian Institute for Health, 2005). Tổng hợp lại, các phát hiện trên cho thấy stress tăng có liên quan đến sử dụng chất ở học sinh và hỗ trợ xã hội nhận thức được — bao gồm từ trong môi trường giáo dục — có thể làm trung gian tích cực cho mối quan hệ này."
)

H('Academic-related stress and sleep — Stress học đường và giấc ngủ', level=2)
pair(
    "Insufficient sleep in adolescents is recognised as a serious health risk by the American Medical Association and the American Academy of Sleep Medicine, who report that many young people do not get enough hours of sleep (Owens, 2014). Stress is a contributing factor to poor sleep in young people (Bernert, Merrill, Braithwaite, Van Orden, & Joiner, 2007; Curcio, Ferrara, & De, 2006). Noland et al., found that 42% of 9–12th Grade students report that stress is an impediment to good sleep, in 384 students surveyed (Noland, Price, Dake, & Telljohann, 2009).",
    "Thiếu ngủ ở vị thành niên được Hiệp hội Y khoa Hoa Kỳ và Học viện Y học Giấc ngủ Hoa Kỳ công nhận là một nguy cơ sức khỏe nghiêm trọng — họ báo cáo rằng nhiều người trẻ không ngủ đủ giờ (Owens, 2014). Stress là yếu tố góp phần gây giấc ngủ kém ở người trẻ (Bernert, Merrill, Braithwaite, Van Orden, & Joiner, 2007; Curcio, Ferrara, & De, 2006). Noland và cs. thấy 42% học sinh lớp 9-12 báo cáo stress là rào cản cho giấc ngủ tốt, trong khảo sát 384 học sinh (Noland, Price, Dake, & Telljohann, 2009)."
)

pair(
    "Self-perceived stress has been shown to result in poorer sleep in female university students from the United States (Lee, Wuertz, Rogers, & Chen, 2013; Wallace, Boynton, & Lytle, 2017), medical students from Saudi Arabia (Almojali, Almalki, Alothman, Masuadi, & Alaqeel, 2017), university students from Portugal (Amaral et al., 2017) and Pakistani medical school students (Waqas, Khan, Sharif, Khalid, & Ali, 2015), demonstrating the cross cultural impacts of stress on sleep quality and quantity tertiary education students. In a study from the United States, over 90% of 9–12th Grade students reported that they have an inadequate number of hours of sleep on most school nights. These young people report that the impact of the loss of sleep is difficulty paying attention,",
    "Stress tự nhận thức được chứng tỏ dẫn đến giấc ngủ kém hơn ở sinh viên đại học nữ tại Mỹ (Lee, Wuertz, Rogers, & Chen, 2013; Wallace, Boynton, & Lytle, 2017), sinh viên y khoa Ả Rập Xê Út (Almojali, Almalki, Alothman, Masuadi, & Alaqeel, 2017), sinh viên đại học Bồ Đào Nha (Amaral và cs., 2017) và sinh viên y khoa Pakistan (Waqas, Khan, Sharif, Khalid, & Ali, 2015) — chứng tỏ tác động xuyên văn hóa của stress lên chất lượng và số lượng giấc ngủ ở sinh viên đại học. Trong một nghiên cứu tại Mỹ, hơn 90% học sinh lớp 9-12 báo cáo họ ngủ không đủ giờ vào hầu hết các đêm trong ngày học. Những thanh niên này báo cáo tác động của thiếu ngủ là khó tập trung,"
)

# ============================================================
# TRANG 6 — SLEEP cont + PHYSICAL HEALTH + ACHIEVEMENT
# ============================================================
page_marker(6)
pair(
    "lower grades, higher stress, and trouble getting along with other people. Some students reported problematic coping strategies such as taking sleeping pills, smoking cigarettes and drinking alcohol to help them sleep (Noland et al., 2009). Sleep quality and quantity has been shown to be closely related to student learning capability and academic performance (Curcio et al., 2006). Loss of sleep is frequently associated with poor learning (Curcio et al., 2006). Therefore, stress-related disruption to sleep quality and quantity is an important factor contributing to poor learning and well-being among students.",
    "điểm thấp hơn, stress cao hơn và khó hòa hợp với người khác. Một số học sinh báo cáo các chiến lược ứng phó có vấn đề như uống thuốc ngủ, hút thuốc lá và uống rượu để dễ ngủ (Noland và cs., 2009). Chất lượng và số lượng giấc ngủ được chứng tỏ có liên quan chặt với năng lực học tập và thành tích học tập (Curcio và cs., 2006). Thiếu ngủ thường liên quan đến học tập kém (Curcio và cs., 2006). Do đó, gián đoạn giấc ngủ liên quan đến stress là yếu tố quan trọng góp phần gây học tập và an sinh kém ở học sinh."
)

H('Academic-related stress and physical health — Stress học đường và sức khỏe thể chất', level=2)
pair(
    "The experience of high levels of academic-related stress increases the risk of young people developing preventable physical health problems later in life. A systematic review of prospective studies found that people who were stressed, such as during examination periods, were less likely to be physically active, the impact of which is associated with a plethora of potentially inter-related poor physical health outcomes (Stults-Kolehmainen & Sinha, 2014). Stress may also lead to the development of non-communicable diseases, including metabolic syndrome, obesity and reduced insulin sensitivity, resulting from unhealthy lifestyle habits and stress system dysregulation (Pervanidou & Chrousos, 2012). Similarly, stress has been shown to be associated with increased appetite (Dallman et al., 1993) and higher body weight (Stephens et al., 1995). Therefore, academic-related stress can contribute to the development of health issues, including chronic non-communicable diseases, due to decreases in physical activity and increases in unhealthy lifestyle habits.",
    "Trải nghiệm stress học đường ở mức cao làm tăng nguy cơ người trẻ phát triển các vấn đề sức khỏe thể chất có thể phòng ngừa trong tương lai. Một tổng quan hệ thống các nghiên cứu thuần tập tiến cứu thấy rằng người bị stress — chẳng hạn trong thời gian thi cử — ít hoạt động thể chất hơn, và tác động của điều này có liên quan đến hàng loạt kết quả sức khỏe thể chất kém có khả năng đan xen lẫn nhau (Stults-Kolehmainen & Sinha, 2014). Stress cũng có thể dẫn đến phát triển các bệnh không lây nhiễm, bao gồm hội chứng chuyển hóa, béo phì và giảm độ nhạy insulin — kết quả của thói quen lối sống không lành mạnh và rối loạn điều hòa hệ thống stress (Pervanidou & Chrousos, 2012). Tương tự, stress được chứng tỏ có liên quan đến tăng cảm giác ngon miệng (Dallman và cs., 1993) và cân nặng cao hơn (Stephens và cs., 1995). Do đó, stress học đường có thể góp phần phát triển các vấn đề sức khỏe — bao gồm bệnh không lây nhiễm mạn tính — do giảm hoạt động thể chất và tăng các thói quen lối sống không lành mạnh."
)

H('Academic-related stress and achievement — Stress học đường và thành tích', level=2)
pair(
    "The World Health Organisation (1996) states that students must be healthy and emotionally secure to fully participate in education (World Health Organisation, 1996). Indeed, the above-mentioned OECD survey reports that anxiety about schoolwork, homework and tests has a negative impact on students' academic performance in science, mathematics and reading. The survey highlights that top-performing girls report that the fear of making mistakes often disrupts their test performance (OECD, 2015). Students in the bottom quarter of academic performance report feeling far more stressed compared to those in the top quarter of academic performance. As many as 63% of students in the bottom quarter of science performance report feeling anxious about tests no matter how well prepared they are, while 46% of students in the top quarter report feeling anxious (OECD, 2015). This demonstrates that higher perceived stress levels are associated with poorer academic performance.",
    "Tổ chức Y tế Thế giới (1996) khẳng định học sinh phải khỏe mạnh và an toàn về cảm xúc để tham gia đầy đủ vào giáo dục (WHO, 1996). Thực tế, khảo sát OECD đã đề cập ở trên báo cáo lo âu về bài vở, bài tập và kiểm tra có tác động tiêu cực đến thành tích học tập của học sinh ở khoa học, toán học và đọc. Khảo sát nhấn mạnh các học sinh nữ thành tích cao nhất báo cáo rằng nỗi sợ mắc lỗi thường làm gián đoạn kết quả thi của họ (OECD, 2015). Học sinh ở nhóm 25% thành tích thấp nhất báo cáo cảm thấy stress nhiều hơn nhiều so với nhóm 25% cao nhất. Tới 63% học sinh nhóm 25% thấp nhất ở môn khoa học báo cáo cảm thấy lo lắng về kiểm tra dù chuẩn bị tốt thế nào, trong khi 46% học sinh nhóm 25% cao nhất báo cáo cảm thấy lo lắng (OECD, 2015). Điều này chứng tỏ mức stress nhận thức cao hơn có liên quan đến thành tích học tập kém hơn."
)

pair(
    "Previous research shows that the experience of positive and negative emotions are directly related to levels of student engagement (Reschly, Huebner, Appleton, & Antaramian, 2008). In 293 students in Grades 7–10 from the United States, the frequency of positive emotions during classes was associated with higher student engagement. Conversely, the frequency of negative emotions was associated with lower engagement (Reschly et al., 2008). This finding is important as engagement in learning is necessary for achievement, as illustrated by the findings of a survey conducted by the National Union of Students. This survey reported that the main factors affecting the tertiary studies of Australian university students aged 17–25 was stress (Rickwood, Telford, O'Sullivan, Crisp, & Magyar, 2016). In an observational study of 456 German undergraduate medical students, higher perceived academic-related stress was found to predict poor academic performance (Kotter, Wagner, Bruheim, & Voltmer, 2017). In another study of 121 medical students from Hong Kong, high self-reported stress levels were similarly related to poorer academic performance (Stewart, Lam, Betson, Wong, & Wong, 1999).",
    "Các nghiên cứu trước cho thấy trải nghiệm các cảm xúc tích cực và tiêu cực có liên quan trực tiếp đến mức độ gắn kết của học sinh (Reschly, Huebner, Appleton, & Antaramian, 2008). Ở 293 học sinh lớp 7-10 tại Mỹ, tần suất cảm xúc tích cực trong giờ học có liên quan đến gắn kết cao hơn của học sinh. Ngược lại, tần suất cảm xúc tiêu cực có liên quan đến gắn kết thấp hơn (Reschly và cs., 2008). Phát hiện này quan trọng vì gắn kết với việc học là cần thiết cho thành tích, như được minh họa bởi khảo sát của Liên hiệp Sinh viên Quốc gia. Khảo sát này báo cáo các yếu tố chính ảnh hưởng tới việc học đại học của sinh viên Úc tuổi 17-25 là stress (Rickwood, Telford, O'Sullivan, Crisp, & Magyar, 2016). Trong nghiên cứu quan sát trên 456 sinh viên y khoa hệ cử nhân ở Đức, stress học đường nhận thức cao hơn được phát hiện dự báo thành tích học tập kém (Kotter, Wagner, Bruheim, & Voltmer, 2017). Trong một nghiên cứu khác trên 121 sinh viên y khoa ở Hong Kong, mức stress tự báo cáo cao tương tự có liên quan đến thành tích học tập kém hơn (Stewart, Lam, Betson, Wong, & Wong, 1999)."
)

pair(
    "The above findings demonstrate that the academic-related stress that secondary and tertiary students experience constitutes a major factor affecting their academic achievement. Students with higher perceived stress are likely to have lower academic achievement.",
    "Các phát hiện trên chứng tỏ stress học đường mà học sinh trung học và sinh viên đại học trải qua tạo nên một yếu tố lớn ảnh hưởng đến thành tích học tập của họ. Học sinh có stress nhận thức cao hơn có khả năng có thành tích học tập thấp hơn."
)

# ============================================================
# TRANG 7 — DROPOUT + DISCUSSION start
# ============================================================
page_marker(7)
H('Academic-related stress and dropout — Stress học đường và bỏ học', level=2)
pair(
    "Academic-related stress and burnout includes exhaustion, depersonalization, cynicism and inefficacy or reduced accomplishment (Walburg, 2014). Academic-related stress is strongly related to decreased student academic motivation (Liu, 2015; Liu & Lu, 2011; Shinto, 1998) and academic disengagement (National Centre on Addiction and Substance Abuse at Columbia University (CASA) United States of America, 2003). The relationship between academic-related stress, motivation and dropout does not appear to be culturally specific, with similar findings shown from a number of international studies (Liu, 2015; Liu & Lu, 2011; Shinto, 1998; Walburg, 2014).",
    "Stress học đường và kiệt sức (burnout) bao gồm kiệt quệ, mất nhân cách hóa (depersonalization), hoài nghi và kém hiệu quả hoặc giảm thành tựu (Walburg, 2014). Stress học đường có liên quan mạnh đến giảm động lực học tập của học sinh (Liu, 2015; Liu & Lu, 2011; Shinto, 1998) và tách rời khỏi học tập (academic disengagement) (Trung tâm Quốc gia về Nghiện và Lạm dụng Chất tại Đại học Columbia (CASA) Hoa Kỳ, 2003). Mối quan hệ giữa stress học đường, động lực và bỏ học có vẻ KHÔNG đặc thù theo văn hóa — phát hiện tương tự được thấy ở nhiều nghiên cứu quốc tế (Liu, 2015; Liu & Lu, 2011; Shinto, 1998; Walburg, 2014)."
)

pair(
    "In 298 Chinese secondary school students, academic-related stress in Grade 10 negatively predicted intrinsic academic motivation and positively predicted lack of motivation in Grade 12. This indicates that decreasing academic-related stress might preserve students' ongoing intrinsic academic motivation (Liu, 2015; Liu & Lu, 2011). Similarly, in 495 Japanese students in junior secondary school, self-reported academic-related stress was found to negatively relate to feelings of self-growth and academic motivation (Shinto, 1998). A recent literature review highlights how stress and burnout can also affect academic achievement by increasing the risk for school dropout (Walburg, 2014). This was particularly true for students who experience more stressful life events of a more severe nature, as well as students who do not seek support from their parents or other family members as well as students from ethnically diverse groups (Hess & Copeland, 2001).",
    "Ở 298 học sinh trung học Trung Quốc, stress học đường ở lớp 10 dự báo NGHỊCH với động lực học nội tại và dự báo THUẬN với thiếu động lực ở lớp 12. Điều này gợi ý giảm stress học đường có thể bảo vệ động lực học nội tại liên tục của học sinh (Liu, 2015; Liu & Lu, 2011). Tương tự, ở 495 học sinh THCS Nhật Bản, stress học đường tự báo cáo có liên quan nghịch với cảm giác tự phát triển và động lực học tập (Shinto, 1998). Một tổng quan tài liệu gần đây nhấn mạnh stress và burnout cũng có thể ảnh hưởng đến thành tích học tập bằng cách tăng nguy cơ bỏ học (Walburg, 2014). Điều này đặc biệt đúng với học sinh trải qua nhiều biến cố căng thẳng nghiêm trọng hơn, cũng như học sinh không tìm kiếm hỗ trợ từ cha mẹ hoặc thành viên gia đình khác, cũng như học sinh từ các nhóm dân tộc đa dạng (Hess & Copeland, 2001)."
)

pair(
    "School dropout is associated with a lifelong reduction in earning capacity and secure employment (Lamb & Huo, 2017). Individuals with lower education levels report having poorer mental health and more illness than those with higher levels of education (Turrell, Stanley, de Looper, & Oldenburg, 2006). Early dropout from school has also been reported to contribute to inter-generational issues including unemployment, poverty and less academic achievement (Black, 2007; Lamb & Huo, 2017; Muir, Family, Maguire, Slack-Smith, & Murray, 2003). Academic achievement and completion of secondary school leads to greater employability, less reliance on social welfare support and a higher likelihood of participation in further education (Noble, Wyatt, McGrath, Roffey, & Rowling, 2008). These outcomes in turn increase the likelihood of sustainable employment, adequate income and self-sufficiency (Noble et al., 2008), which can save Governments hundreds of millions of dollars every year (Lamb & Huo, 2017).",
    "Bỏ học có liên quan đến giảm năng lực kiếm tiền và việc làm ổn định suốt đời (Lamb & Huo, 2017). Cá nhân có trình độ học vấn thấp hơn báo cáo sức khỏe tâm thần kém hơn và nhiều bệnh hơn so với người có trình độ cao hơn (Turrell, Stanley, de Looper, & Oldenburg, 2006). Bỏ học sớm cũng được báo cáo góp phần gây các vấn đề liên thế hệ — bao gồm thất nghiệp, nghèo đói và thành tích học tập thấp hơn (Black, 2007; Lamb & Huo, 2017; Muir, Family, Maguire, Slack-Smith, & Murray, 2003). Thành tích học tập và hoàn thành trung học dẫn đến khả năng có việc làm cao hơn, ít phụ thuộc trợ cấp xã hội hơn và khả năng cao hơn tham gia vào giáo dục bậc cao hơn (Noble, Wyatt, McGrath, Roffey, & Rowling, 2008). Các kết quả này lần lượt tăng khả năng có việc làm bền vững, thu nhập đủ và tự chủ (Noble và cs., 2008) — có thể tiết kiệm hàng trăm triệu đô-la cho các Chính phủ mỗi năm (Lamb & Huo, 2017)."
)

H('Discussion — Thảo luận', level=2)
pair(
    "The current narrative review highlights that students commonly report high levels of academic-related stress, cross-culturally. The academic-related stress experienced by secondary and tertiary students' impacts their mental and physical health and leads to a range of academic problems. Good stress-management skills have the potential to benefit young people in an ongoing manner throughout their lives, given that many long-term health-related behaviours and patterns, both positive and negative, are established during adolescence and early adulthood (Sawyer et al., 2012). Therefore, providing opportunities to improve young people's academic stress-related coping abilities during this highly stressful, crucial period of development is an important target (OECD, 2015).",
    "Bài tổng quan tường thuật này nhấn mạnh rằng học sinh thường báo cáo mức stress học đường cao, xuyên văn hóa. Stress học đường mà học sinh trung học và sinh viên đại học trải qua ảnh hưởng đến sức khỏe tâm thần và thể chất của họ, đồng thời dẫn đến hàng loạt vấn đề học tập. Kỹ năng quản lý stress tốt có tiềm năng mang lại lợi ích cho người trẻ một cách liên tục suốt cuộc đời họ — vì nhiều hành vi và mẫu hình liên quan đến sức khỏe lâu dài (cả tích cực và tiêu cực) được hình thành trong thời vị thành niên và đầu trưởng thành (Sawyer và cs., 2012). Do đó, cung cấp cơ hội cải thiện khả năng ứng phó stress học đường của người trẻ trong giai đoạn phát triển then chốt và đầy stress này là mục tiêu quan trọng (OECD, 2015)."
)

pair(
    "The OECD highlights that education settings are places where young people develop many of the social and emotional skills needed to become resilient and thrive (OECD, 2015). Therefore, education settings can work to improve student academic related stress through the provision of programmes shown to decrease stress and increase stress management and coping. Discussion regarding the efficacy of particular school based stress management programmes to teach students to cope with stress is beyond the scope of the current review. It worth noting, however, that education-based initiatives that focus on increasing students skills and ability to cope with stress have been previously demonstrated to directly and positively influence educational achievement and decrease health risks (Hanson & Austin, 2002; Perry et al., 2017; Weare & Gray, 2003). For example, a meta-analysis of 19 randomised controlled",
    "OECD nhấn mạnh môi trường giáo dục là nơi người trẻ phát triển nhiều kỹ năng xã hội và cảm xúc cần thiết để trở nên kiên cường (resilient) và phát triển mạnh (OECD, 2015). Do đó, môi trường giáo dục có thể nỗ lực cải thiện stress học đường của học sinh thông qua cung cấp các chương trình được chứng tỏ làm giảm stress và tăng quản lý stress và ứng phó. Thảo luận về hiệu quả của các chương trình quản lý stress cụ thể tại trường để dạy học sinh ứng phó với stress nằm ngoài phạm vi của bài tổng quan này. Tuy nhiên đáng lưu ý rằng các sáng kiến dựa vào giáo dục tập trung vào tăng kỹ năng và khả năng ứng phó với stress của học sinh đã được chứng tỏ ảnh hưởng trực tiếp và tích cực đến thành tích giáo dục và giảm nguy cơ sức khỏe (Hanson & Austin, 2002; Perry và cs., 2017; Weare & Gray, 2003). Ví dụ, một phân tích tổng hợp 19 nghiên cứu ngẫu nhiên có đối chứng"
)

# ============================================================
# TRANG 8 — DISCUSSION cont + CONCLUSION + DISCLOSURE + REFERENCES start
# ============================================================
page_marker(8)
pair(
    "trials or quasi-experimental studies found that school programmes targeting stress management or coping skills reduced stress symptoms and improved coping skills among students (Kraag, Zeegers, Kok, Hosman, & Abu-Saad, 2006). Schools provide access to a large number of young people, across a diverse range of backgrounds, during a formative developmental period (Sawyer et al., 2012). As such, even if modestly effective, the population level implementation of stress management and coping skills programmes would help young people to develop healthy coping strategies in order to deal with the inevitable stressors of life. Understanding and addressing the barriers and enablers to implementation of stress management programmes in schools would support the development of effective implementation strategies (Albers & Pattuwage, 2017; Domitrovich et al., 2008), resulting in significant health, economic and social benefits for large numbers of young people, their families and the community.",
    "hoặc bán thực nghiệm thấy các chương trình ở trường nhắm vào quản lý stress hoặc kỹ năng ứng phó làm giảm các triệu chứng stress và cải thiện kỹ năng ứng phó ở học sinh (Kraag, Zeegers, Kok, Hosman, & Abu-Saad, 2006). Trường học cung cấp khả năng tiếp cận một lượng lớn người trẻ thuộc nhiều nền tảng đa dạng, trong giai đoạn phát triển hình thành (Sawyer và cs., 2012). Như vậy, ngay cả khi chỉ có hiệu quả ở mức trung bình, việc triển khai các chương trình quản lý stress và kỹ năng ứng phó ở cấp dân số sẽ giúp người trẻ phát triển các chiến lược ứng phó lành mạnh để đối phó với các stressor không thể tránh khỏi của cuộc sống. Hiểu và giải quyết các rào cản và yếu tố tạo thuận lợi cho triển khai chương trình quản lý stress ở trường sẽ hỗ trợ phát triển các chiến lược triển khai hiệu quả (Albers & Pattuwage, 2017; Domitrovich và cs., 2008) — dẫn đến lợi ích sức khỏe, kinh tế và xã hội đáng kể cho nhiều người trẻ, gia đình họ và cộng đồng."
)

pair(
    "A strength of the current review is that we have discussed studies from many countries, indicating that the academic-related stress experienced by students in education is cross-cultural and wide spread and is of international concern. We reviewed studies that demonstrated a range of negative effects of academic-related stress, highlighting the potential broad spectrum of benefits that may result from the implementation of stress-management interventions. A limitation of the current study is that we have not delineated between studies that have assessed the impact of academic-related stress during different phases of secondary and tertiary education. It is more than likely that the needs and therefore the most beneficial coping strategies may vary throughout the life span. Therefore, the most appropriate stress-management education approaches may differ between the early high school and tertiary education years.",
    "Một điểm mạnh của bài tổng quan hiện tại là chúng tôi đã thảo luận các nghiên cứu từ nhiều quốc gia, cho thấy stress học đường mà học sinh trải nghiệm trong giáo dục là xuyên văn hóa, lan rộng và là mối quan ngại quốc tế. Chúng tôi xem xét các nghiên cứu chứng minh hàng loạt tác động tiêu cực của stress học đường — nhấn mạnh phổ rộng tiềm năng của các lợi ích có thể có từ việc triển khai các can thiệp quản lý stress. Một hạn chế của nghiên cứu hiện tại là chúng tôi chưa phân biệt giữa các nghiên cứu đánh giá tác động của stress học đường trong các giai đoạn khác nhau của trung học và đại học. Khả năng cao là nhu cầu — và do đó các chiến lược ứng phó có lợi nhất — có thể thay đổi theo độ tuổi đời. Do đó, các phương pháp giáo dục quản lý stress phù hợp nhất có thể khác nhau giữa những năm đầu THPT và những năm đại học."
)

H('Conclusion — Kết luận', level=2)
pair(
    "This narrative review highlights that academic-related stress is a major concern for secondary and tertiary students. The ongoing stress relating to education has demonstrated negative impact on students' learning capacity, academic performance, education and employment attainment, sleep quality and quantity, physical health, mental health and substance use outcomes. Increasing students' stress-management skills and abilities is an important target for change.",
    "Bài tổng quan tường thuật này nhấn mạnh rằng stress học đường là mối quan ngại lớn đối với học sinh trung học và sinh viên đại học. Stress liên quan đến giáo dục liên tục đã được chứng tỏ có tác động tiêu cực đến năng lực học tập, thành tích học tập, đạt được giáo dục và việc làm, chất lượng và số lượng giấc ngủ, sức khỏe thể chất, sức khỏe tâm thần và kết quả sử dụng chất của học sinh. Tăng kỹ năng và khả năng quản lý stress của học sinh là mục tiêu quan trọng để thay đổi."
)

H('Disclosure / Funding / ORCID', level=3)
pair(
    "Disclosure statement: No potential conflict of interest was reported by the authors. Funding: The authors received no financial support for the research, authorship and/or publication of this article. ORCID: Michaela C. Pascoe http://orcid.org/0000-0002-3831-5660",
    "Tuyên bố công khai: Các tác giả không báo cáo xung đột lợi ích tiềm tàng nào. Tài trợ: Các tác giả không nhận hỗ trợ tài chính nào cho nghiên cứu, soạn thảo và/hoặc xuất bản bài báo này. ORCID: Michaela C. Pascoe http://orcid.org/0000-0002-3831-5660"
)

# ============================================================
# REFERENCES — bilingual list
# ============================================================
H('References — Tài liệu tham khảo (song ngữ)', level=2)
note('Theo Nguyên tắc dịch v2 (Nguyên tắc 2): tài liệu tham khảo giữ nguyên văn tiếng Anh để có thể tra cứu lại; phần dòng → in nghiêng phía dưới là CHÚ DỊCH ngắn gọn tiêu đề tiếng Việt, không thay thế bản gốc.', color=GREEN)

REFS = [
    ("Albers, B., & Pattuwage, L. (2017). Implementation in education: Findings from a scoping review. Melbourne: Evidence for Learning. doi: 10.13140/RG.2.2.29187.40483",
     "Triển khai trong giáo dục: Phát hiện từ tổng quan khảo lược."),
    ("Almojali, A. I., Almalki, S. A., Alothman, A. S., Masuadi, E. M., & Alaqeel, M. K. (2017). The prevalence and association of stress with sleep quality among medical students. Journal of Epidemiology and Global Health, 7(3), 169–174.",
     "Tỷ lệ stress và mối liên hệ với chất lượng giấc ngủ ở sinh viên y khoa."),
    ("Amaral, A. P., Soares, M. J., Pinto, A. M., Pereira, A. T., Madeira, N., Bos, S. C., … Macedo, A. (2017). Sleep difficulties in college students: The role of stress, affect and cognitive processes. Psychiatry Research, 260, 331–337.",
     "Khó ngủ ở sinh viên đại học: vai trò của stress, cảm xúc và quá trình nhận thức."),
    ("Bayram, N., & Bilgel, N. (2008). The prevalence and socio-demographic correlations of depression, anxiety and stress among a group of university students. Social Psychiatry and Psychiatric Epidemiology, 43(8), 667–672.",
     "Tỷ lệ và tương quan nhân khẩu xã hội của trầm cảm, lo âu và stress ở sinh viên đại học."),
    ("Bernal-Morales, B., Rodríguez-Landa, J. F., & Pulido-Criollo, F. (2015). Impact of anxiety and depression symptoms on scholar performance in high school and university students. London, UK: IntechOpen.",
     "Tác động của triệu chứng lo âu và trầm cảm lên thành tích học tập ở học sinh THPT và sinh viên."),
    ("Bernert, R. A., Merrill, K. A., Braithwaite, S. R., Van Orden, K. A., & Joiner, T. E., Jr. (2007). Family life stress and insomnia symptoms in a prospective evaluation of young adults. Journal of Family Psychology, 21(1), 58–66.",
     "Stress đời sống gia đình và triệu chứng mất ngủ ở thanh niên — đánh giá tiến cứu."),
    ("Black, R. (2007). Crossing the bridge: overcoming entrenched disadvantage through student-centred learning. Melbourne: Education Foundation Australia.",
     "Vượt qua bất lợi ăn sâu qua học tập lấy học sinh làm trung tâm."),
    ("Boulton, M., & O'Connell, K. A. (2017). Nursing students' perceived faculty support, stress, and substance misuse. The Journal of Nursing Education, 56(7), 404–411.",
     "Hỗ trợ từ giảng viên (nhận thức), stress và lạm dụng chất ở sinh viên điều dưỡng."),
    ("Canadian Institute for Health. (2005). Improving the health of young Canadians. Ottawa, Ontario: Canadian Institute for Health Information.",
     "Cải thiện sức khỏe người trẻ Canada (Viện Y tế Canada)."),
    ("Carter, J. S., Garber, J., Ciesla, J. A., & Cole, D. A. (2006). Modeling relations between hassles and internalizing and externalizing symptoms in adolescents: A four-year prospective study. Journal of Abnormal Psychology, 115(3), 428–442.",
     "Mô hình hóa quan hệ giữa các phiền toái và triệu chứng hướng nội/ngoại ở vị thành niên — thuần tập 4 năm."),
    ("Chapell, M. S., Blanding, Z. B., Silverstein, M. E., Takahashi, M., Newman, B., Gubi, A., & McCann, N. (2005). Test anxiety and academic performance in undergraduate and graduate students. Journal of Educational Psychology, 97(2), 268–274.",
     "Lo âu kiểm tra và thành tích học tập ở sinh viên cử nhân và sau đại học."),
    ("Curcio, G., Ferrara, M., & De, G. L. (2006). Sleep loss, learning capacity and academic performance. Sleep Medicine Reviews, 10(5), 323–337.",
     "Thiếu ngủ, năng lực học tập và thành tích học tập."),
    ("Dallman, M. F., Strack, A. M., Akana, S. F., Bradbury, M. J., Hanson, E. S., Scribner, K. A., & Smith, M. (1993). Feast and famine: Critical role of glucocorticoids with insulin in daily energy flow. Frontiers in Neuroendocrinology, 14(4), 303–347.",
     "Vai trò then chốt của glucocorticoid với insulin trong dòng năng lượng hàng ngày."),
    ("Dantzer, R. (2012). Depression and inflammation: An intricate relationship. Biological Psychiatry, 71(1), 4–5.",
     "Trầm cảm và viêm: một mối quan hệ phức tạp."),
    ("Dantzer, R., O'Connor, J. C., Lawson, M. A., & Kelley, K. W. (2011). Inflammation-associated depression: From serotonin to kynurenine. Psychoneuroendocrinology, 36(3), 426–436.",
     "Trầm cảm liên quan đến viêm: từ serotonin đến kynurenine."),
    ("Domitrovich, C. E., Bradshaw, C. P., Poduska, J. M., Hoagwood, K., Buckley, J. A., Olin, S., … Ialongo, N. S. (2008). Maximizing the implementation quality of evidence-based preventive interventions in schools. Advances in School Mental Health Promotion, 1(3), 6–28.",
     "Tối đa hóa chất lượng triển khai các can thiệp phòng ngừa dựa trên bằng chứng tại trường."),
    ("Eisenberg, D., Gollust, S. E., Golberstein, E., & Hefner, J. L. (2007). Prevalence and correlates of depression, anxiety, and suicidality among university students. The American Journal of Orthopsychiatry, 77(4), 534–542.",
     "Tỷ lệ và yếu tố liên quan của trầm cảm, lo âu và ý định tự tử ở sinh viên đại học."),
    ("Fergusson, D. M., Boden, J. M., & Horwood, L. J. (2007). Recurrence of major depression in adolescence and early adulthood, and later mental health, educational and economic outcomes. The British Journal of Psychiatry, 191, 335–342.",
     "Tái phát trầm cảm chính ở vị thành niên và đầu trưởng thành — kết quả sức khỏe tâm thần, giáo dục và kinh tế về sau."),
    ("Fröjd, S. A., Nissinen, E. S., Pelkonen, M. U. I., Marttunen, M. J., Koivisto, A.-M., & Kaltiala-Heino, R. (2008). Depression and school performance in middle adolescent boys and girls. Journal of Adolescence, 31(4), 485–498.",
     "Trầm cảm và thành tích học tập ở nam và nữ vị thành niên giữa."),
    ("Hanson, T. L., & Austin, G. A. (2002). Health risks, resilience, and the academic performance index (California healthy kids survey factsheet 1). Los Alamitos, CA: WestEd.",
     "Nguy cơ sức khỏe, kiên cường và chỉ số thành tích học tập (khảo sát trẻ khỏe California)."),
    ("Hess, R. S., & Copeland, E. P. (2001). Students' stress, coping strategies, and school completion: A longitudinal perspective. School Psychology Quarterly, 16(4), 389.",
     "Stress học sinh, chiến lược ứng phó và hoàn thành trường — góc nhìn thuần tập."),
    ("Humensky, J., Kuwabara, S. A., Fogel, J., Wells, C., Goodwin, B., & Van Voorhees, B. W. (2010). Adolescents with depressive symptoms and their challenges with learning in school. The Journal of School Nursing, 26(5), 377–392.",
     "Vị thành niên có triệu chứng trầm cảm và thách thức học tập tại trường."),
    ("Hysenbegasi, A., Hass, S. L., & Rowland, C. R. (2005). The impact of depression on the academic productivity of university students. The Journal of Mental Health Policy and Economics, 8(3), 145–151.",
     "Tác động của trầm cảm lên năng suất học tập của sinh viên đại học."),
    ("Ibrahim, A. K., Kelly, S. J., Adams, C. E., & Glazebrook, C. (2013). A systematic review of studies of depression prevalence in university students. Journal of Psychiatric Research, 47(3), 391–400.",
     "Tổng quan hệ thống các nghiên cứu về tỷ lệ trầm cảm ở sinh viên đại học."),
    ("Jonsson, U., Bohman, H., Hjern, A., von Knorring, L., Olsson, G., & von Knorring, A. L. (2010). Subsequent higher education after adolescent depression: A 15-year follow-up register study. European Psychiatry, 25(7), 396–401. doi:10.1016/j.eurpsy.2010.01.016",
     "Học đại học sau trầm cảm vị thành niên — nghiên cứu sổ đăng ký theo dõi 15 năm."),
    ("Kendler, K. S., Kessler, R. C., Walters, E. E., MacLean, C., Neale, M. C., Heath, A. C., & Eaves, L. J. (1995). Stressful life events, genetic liability, and onset of an episode of major depression in women. The American Journal of Psychiatry, 152(6), 833–842.",
     "Biến cố căng thẳng đời sống, khuynh hướng di truyền và khởi phát giai đoạn trầm cảm chính ở phụ nữ."),
    ("Kessler, R. C. (1997). The effects of stressful life events on depression. Annual Review of Psychology, 48, 191–214.",
     "Ảnh hưởng của các biến cố căng thẳng đời sống lên trầm cảm."),
    ("Kessler, R. C. (2012). The costs of depression. The Psychiatric Clinics of North America, 35(1), 1–14.",
     "Chi phí của trầm cảm."),
    ("Kotter, T., Wagner, J., Bruheim, L., & Voltmer, E. (2017). Perceived Medical School stress of undergraduate medical students predicts academic performance: An observational study. BMC Medical Education, 17(1), 256.",
     "Stress trường y nhận thức của sinh viên y khoa cử nhân dự báo thành tích học tập — nghiên cứu quan sát."),
    ("Kraag, G., Zeegers, M. P., Kok, G., Hosman, C., & Abu-Saad, H. H. (2006). School programs targeting stress management in children and adolescents: A meta-analysis. Journal of School Psychology, 44(6), 449–472.",
     "Các chương trình trường học nhắm vào quản lý stress ở trẻ em và vị thành niên — phân tích tổng hợp."),
    ("Lamb, S., & Huo, S. (2017). Counting the costs of lost opportunity in Australian education. Melbourne, Australia: Mitchell Institute.",
     "Tính toán chi phí cơ hội bị mất trong giáo dục Úc."),
    ("Lee, S. Y., Wuertz, C., Rogers, R., & Chen, Y. P. (2013). Stress and sleep disturbances in female college students. American Journal of Health Behavior, 37(6), 851–858.",
     "Stress và rối loạn giấc ngủ ở sinh viên đại học nữ."),
    ("Leonard, N. R., Gwadz, M. V., Ritchie, A., Linick, J. L., Cleland, C. M., Elliott, L., & Grethel, M. (2015). A multi-method exploratory study of stress, coping, and substance use among high school youth in private schools. Frontiers in Psychology, 6, 1028.",
     "Nghiên cứu khám phá đa phương pháp về stress, ứng phó và sử dụng chất ở thanh thiếu niên trường tư."),
    ("Lewinsohn, P. M., Allen, N. B., Seeley, J. R., & Gotlib, I. H. (1999). First onset versus recurrence of depression: Differential processes of psychosocial risk. Journal of Abnormal Psychology, 108, 483–489.",
     "Khởi phát đầu tiên so với tái phát trầm cảm: các quá trình rủi ro tâm lý-xã hội khác nhau."),
    ("Liu, Y. Y. (2015). The longitudinal relationship between Chinese high school students' academic stress and academic motivation. Learning and Individual Differences, 38, 123–126.",
     "Mối quan hệ thuần tập giữa stress học đường và động lực học tập ở học sinh THPT Trung Quốc."),
    ("Liu, Y. Y., & Lu, Z. H. (2011). The Chinese high school student's stress in the school and academic achievement. Educational Psychology, 31(1), 27–35.",
     "Stress học đường và thành tích học tập của học sinh THPT Trung Quốc."),
    ("Maes, M. (2008). The cytokine hypothesis of depression: Inflammation, oxidative & nitrosative stress (IO&NS) and leaky gut as new targets for adjunctive treatments in depression. Neuroendocrinology Letters, 29(3), 287–291.",
     "Giả thuyết cytokine về trầm cảm: viêm, stress oxy hóa và rò rỉ ruột là mục tiêu mới."),
    ("McArdle, J., Hamagami, F., Chang, J. Y., & Hishinuma, E. S. (2014). Longitudinal dynamic analyses of depression and academic achievement in the Hawaiian high schools health survey. Structural Equation Modeling, 21(4), 608–629.",
     "Phân tích động thuần tập về trầm cảm và thành tích học tập ở khảo sát trường THPT Hawaii."),
    ("Moylan, S., Maes, M., Wray, N. R., & Berk, M. (2013). The neuroprogressive nature of major depressive disorder: Pathways to disease evolution and resistance, and therapeutic implications. Molecular Psychiatry, 18(5), 595–606.",
     "Bản chất tiến triển thần kinh của rối loạn trầm cảm chính: con đường tiến hóa bệnh và đề kháng, hệ quả điều trị."),
    ("Muir, K., Family, S., Maguire, A., Slack-Smith, D., & Murray, M. (2003). Youth unemployment in Australia: A contextual, governmental and organisational perspective. Camperdown: The Smith Family.",
     "Thất nghiệp thanh niên ở Úc: góc nhìn bối cảnh, chính phủ và tổ chức."),
    ("National Centre on Addiction and Substance Abuse at Columbia University (CASA) USA. (2003). Depression, substance abuse and college student engagement: A review of the literature. New York: CASA.",
     "Trầm cảm, lạm dụng chất và gắn kết của sinh viên đại học — tổng quan tài liệu."),
    ("Noble, T., Wyatt, T., McGrath, H., Roffey, S., & Rowling, L. (2008). Scoping study into approaches to student wellbeing - final report. Brisbane, Australia.",
     "Nghiên cứu khảo lược về các phương pháp tiếp cận an sinh học sinh — báo cáo cuối."),
    ("Noland, H., Price, J. H., Dake, J., & Telljohann, S. K. (2009). Adolescents' sleep behaviors and perceptions of sleep. The Journal of School Health, 79(5), 224–230.",
     "Hành vi và nhận thức về giấc ngủ ở vị thành niên."),
    ("OECD. (2015). PISA 2015 Results (Volume III). Paris, France.",
     "Kết quả PISA 2015 — Tập III (OECD)."),
    ("OECD. (2017). PISA 2015 Results (Volume III). Paris, France.",
     "Kết quả PISA 2015 — Tập III (OECD, bản 2017)."),
    ("Owens, J. (2014). Insufficient sleep in adolescents and young adults: An update on causes and consequences. Pediatrics.",
     "Thiếu ngủ ở vị thành niên và thanh niên: cập nhật nguyên nhân và hệ quả."),
    ("Ozen, N. S., Ercan, I., Irgil, E., & Sigirli, D. (2010). Anxiety prevalence and affecting factors among university students. Asia-Pacific Journal of Public Health, 22(1), 127–133.",
     "Tỷ lệ lo âu và các yếu tố ảnh hưởng ở sinh viên đại học."),
    ("Perry, Y., Werner-Seidler, A., Calear, A., Mackinnon, A., King, C., Scott, J., … Batterham, P. J. (2017). Preventing depression in final year secondary students: school-based randomized controlled trial. Journal of Medical Internet Research, 19(11), e369.",
     "Phòng ngừa trầm cảm ở học sinh năm cuối THPT — RCT dựa trên trường."),
    ("Pervanidou, P., & Chrousos, G. P. (2012). Metabolic consequences of stress during childhood and adolescence. Metabolism, 61(5), 611–619.",
     "Hệ quả chuyển hóa của stress trong tuổi thơ và vị thành niên."),
    ("Reschly, A. L., Huebner, E. S., Appleton, J. J., & Antaramian, S. (2008). Engagement as flourishing: The contribution of positive emotions and coping to adolescents' engagement at school and with learning. Psychology in the Schools, 45(5), 419–431.",
     "Gắn kết như sự phát triển mạnh: đóng góp của cảm xúc tích cực và ứng phó."),
    ("Ribeiro, Í. J. S., Pereira, R., Freire, I. V., de Oliveira, B. G., Casotti, C. A., & Boery, E. N. (2017). Stress and quality of life among university students: A systematic literature review. Health Professions Education.",
     "Stress và chất lượng cuộc sống ở sinh viên đại học — tổng quan tài liệu hệ thống."),
    ("Rickwood, D., Telford, N., O'Sullivan, S., Crisp, D., & Magyar, R. (2016). National tertiary student wellbeing survey 2016.",
     "Khảo sát an sinh sinh viên đại học toàn quốc 2016 (Úc)."),
    ("Robotham, D., & Julian, C. (2006). Stress and the higher education student: A critical review of the literature. Journal of Further and Higher Education, 30(2), 107–117.",
     "Stress và sinh viên giáo dục bậc cao — tổng quan phê phán tài liệu."),
    ("Ross, C. E., & Mirowsky, J. (2006). Sex differences in the effect of education on depression: Resource multiplication or resource substitution? Social Science & Medicine, 63(5), 1400–1413.",
     "Khác biệt giới trong ảnh hưởng của giáo dục lên trầm cảm: nhân lên hay thay thế nguồn lực?"),
    ("Sawyer, S. M., Afifi, R. A., Bearinger, L. H., Blakemore, S. J., Dick, B., Ezeh, A. C., & Patton, G. C. (2012). Adolescence: A foundation for future health. The Lancet, 379, 1630–1640.",
     "Vị thành niên: nền tảng cho sức khỏe tương lai (loạt bài Lancet)."),
    ("Shinto, T. (1998). Effects of academic stressors and coping strategies on stress responses, feeling of self-growth and motivation in junior high school students. Japanese Journal of Educational Psychology, 46(4), 442–451.",
     "Ảnh hưởng của stressor học đường và chiến lược ứng phó lên phản ứng stress, cảm giác tự phát triển và động lực ở học sinh THCS Nhật."),
    ("Stephens, T. W., Basinski, M., Bristow, P. K., Bue-Valleskey, J. M., Burgett, S. G., Craft, L., … Heiman, M. (1995). The role of neuropeptide Y in the antiobesity action of the obese gene product. Nature, 377(6549), 530–532.",
     "Vai trò của neuropeptide Y trong tác dụng chống béo phì của sản phẩm gen béo phì."),
    ("Stewart, S. M., Lam, T. H., Betson, C. L., Wong, C. M., & Wong, A. M. (1999). A prospective analysis of stress and academic performance in the first two years of medical school. Medical Education, 33(4), 243–250.",
     "Phân tích tiến cứu về stress và thành tích học tập trong 2 năm đầu trường y."),
    ("Stults-Kolehmainen, M. A., & Sinha, R. (2014). The effects of stress on physical activity and exercise. Sports Medicine, 44(1), 81–121.",
     "Ảnh hưởng của stress lên hoạt động thể chất và tập luyện."),
    ("Tountas, Y., & Dimitrakaki, C. (2006). Health education for youth. Pediatric Endocrinology Reviews: PER, 3(Suppl 1), 222–225.",
     "Giáo dục sức khỏe cho thanh niên."),
    ("Turrell, G., Stanley, L., de Looper, M., & Oldenburg, B. (2006). Health inequalities in Australia: Morbidity, health factors, risk factors and health service use. AIHW Cat. No. PHE 72. Canberra: AIHW.",
     "Bất bình đẳng sức khỏe ở Úc: bệnh tật, yếu tố sức khỏe, nguy cơ và sử dụng dịch vụ y tế."),
    ("UNESCO. (2012). International standard classification of education (ISCED) 2011. Montreal, Quebec.",
     "Phân loại tiêu chuẩn quốc tế về giáo dục (ISCED) 2011."),
    ("Walburg, V. (2014). Burnout among high school students: A literature review. Children and Youth Services Review, 42, 28–33.",
     "Burnout ở học sinh THPT — tổng quan tài liệu."),
    ("Wallace, D. D., Boynton, M. H., & Lytle, L. A. (2017). Multilevel analysis exploring the links between stress, depression, and sleep problems among two-year college students. Journal of American College Health, 65(3), 187–196.",
     "Phân tích đa cấp khám phá liên kết giữa stress, trầm cảm và vấn đề giấc ngủ ở sinh viên cao đẳng 2 năm."),
    ("Waqas, A., Khan, S., Sharif, W., Khalid, U., & Ali, A. (2015). Association of academic stress with sleeping difficulties in medical students of a Pakistani medical school: A cross sectional survey. PeerJ, 3, e840.",
     "Liên hệ giữa stress học đường với khó ngủ ở sinh viên y khoa một trường y Pakistan — khảo sát cắt ngang."),
    ("Weare, K., & Gray, G. (2003). What works in developing children's emotional and social competence and wellbeing? Southampton, UK: National Children's Bureau.",
     "Cái gì hiệu quả trong phát triển năng lực cảm xúc-xã hội và an sinh của trẻ?"),
    ("Wills, T. A. (1986). Stress and coping in early adolescence: Relationships to substance use in urban school samples. Health Psychology, 5(6), 503–529.",
     "Stress và ứng phó ở đầu vị thành niên: quan hệ với sử dụng chất ở trường đô thị."),
    ("World Health Organisation. (1996). Health promoting schools. Manilla, Spain.",
     "Trường học thúc đẩy sức khỏe (WHO 1996)."),
    ("World Health Organisation. (2004). Young people's health in context. Health behaviour in school-aged children study: International report from the 2001-2 survey. Denmark.",
     "Sức khỏe người trẻ trong bối cảnh — nghiên cứu hành vi sức khỏe ở trẻ tuổi đi học, báo cáo quốc tế khảo sát 2001-2 (WHO)."),
]

for en_ref, vn_ref in REFS:
    ref_pair(en_ref, vn_ref)

# ============================================================
# PHẢN BIỆN CHI TIẾT (em viết)
# ============================================================
H('PHẢN BIỆN CHI TIẾT — Critique (do dịch giả viết theo Nguyên tắc 9)', level=2, color=RED)

note('1. KHÔNG PHẢI systematic review. Chỉ 1 tác giả tìm trên 2 database (PubMed, Google Scholar), KHÔNG có PRISMA flow, KHÔNG có inclusion/exclusion criteria rõ ràng, KHÔNG có risk-of-bias assessment. Người đọc PHẢI hiểu đây là narrative review — không thể dùng làm "evidence summary" cho meta-analysis.')

note('2. Thiếu dữ liệu Châu Á — Đông Nam Á — Việt Nam. Toàn bộ ~70 references có 6 nghiên cứu Châu Á (Liu Trung Quốc x2, Shinto Nhật, Stewart Hong Kong, Almojali Saudi, Waqas Pakistan), không có Việt Nam, Indonesia, Philippines, Thái Lan. Khi áp dụng vào VN, phải kết hợp với VN013 (Tô Thị Hồng), VN016 (Nguyễn Cao Minh), VN002 (VNAMHS).')

note('3. Trộn lẫn stress học đường với stress chung. Nhiều nghiên cứu được trích dùng PSS (Perceived Stress Scale) chứ không phải thang đo academic stress chuyên biệt (ESSA, ASS). Pascoe không phân biệt rõ → đọc giả dễ nhầm.')

note('4. Số liệu trích dẫn KHÔNG VERIFY ngược nguồn gốc. Ví dụ "prevalence of anxiety as high as 35% in tertiary students" trích Bayram & Bilgel 2008 — nhưng Bayram đo bằng DASS (general anxiety), không phải SAD/social anxiety. Khi dùng số 35% phải đọc trực tiếp Bayram để biết bối cảnh.')

note('5. Mối quan hệ U-NGƯỢC giữa stress và performance (low/moderate/high) KHÔNG được Pascoe nhấn mạnh — chỉ ngầm hiểu. Đây là điểm thiếu sót: stress thấp không phải lúc nào cũng tốt (động lực thấp), stress vừa có thể tối ưu, stress cao mới hại. Pascoe trình bày quan hệ tuyến tính (more stress = worse outcome) không chính xác hoàn toàn.')

note('6. Substance use ở phương Tây (rượu, ma túy, thuốc lá) ≠ substance use ở VN. Ở VN, "chất" thay thế cho HS THCS/THPT là game online, mạng xã hội, tiktok, đồ ăn vặt. Khi áp dụng khung Pascoe vào VN cần localize.')

note('7. Đề xuất can thiệp (mindfulness, CBT, whole-school) đúng hướng nhưng KHÔNG cung cấp evidence cụ thể về effect size. Pascoe nhắc Kraag 2006 meta-analysis nhưng không trích pooled effect size. Khi viết đề cương, em (HLC) phải tìm Kraag 2006 trực tiếp — không trích chéo qua Pascoe.')

# ============================================================
# Ghi chú quy trình
# ============================================================
H('Ghi chú quy trình', level=2)
p = doc.add_paragraph()
r = p.add_run('Ngày dịch: 2026-05-01. Áp dụng Nguyên tắc dịch v2 (10 nguyên tắc): 1-1 paragraph EN/VN, references song ngữ (EN nguyên + VN gloss), markers trang gốc, phản biện đỏ cuối doc. Bài KHÔNG có hình/bảng nên không có phần ảnh+caption. Nguồn: PDF Victoria University Repository (open access CC BY 4.0).')
r.italic = True; r.font.size = Pt(10); r.font.color.rgb = GREY

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
