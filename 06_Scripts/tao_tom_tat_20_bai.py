# -*- coding: utf-8 -*-
"""
Tao 20 ban tom tat moi cho cac bai 39-58 — luu vao Tom-tat-tung-bai/
Theo format mau '00_Mẫu tóm tắt bài 1.docx':
- Tieu de + thong tin thu muc
- Phuong phap nghien cuu
- Ket qua nghien cuu (co bang neu can)
- Nhan xet / phan bien
- Huong nghien cuu tiep theo
- Danh gia chat luong
"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0  # cm

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')):
        tcW.remove(old)
    tcW.append(w)

def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout')
    layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblPr.append(layout)
    tblGrid = t._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        t._tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tblGrid)

def make_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3); s.right_margin = Cm(2)
    return doc

def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
    return t

def save(doc, name):
    p = os.path.join(OUT_DIR, name)
    doc.save(p)
    d = Document(p)
    chars = sum(len(x.text) for x in d.paragraphs)
    print(f'  {name}: {chars} chars, {len(d.tables)} tables')

if __name__ != '__main__':
    import sys as _sys
    _sys.exit(0)

# ============================================================
# BÀI 39 — Nguyen LX 2023 (Medicine, VN COVID GAD-7 n=5.730)
# ============================================================
print('\n[1/20] Bài 39: Nguyen LX 2023 — VN COVID Medicine')
doc = make_doc()
H(doc, 'Tóm tắt bài VN23 — Lo âu trong đại dịch COVID-19 ở sinh viên đại học Việt Nam', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Lo âu ở sinh viên đại học Việt Nam trong đại dịch COVID-19: Khảo sát quốc gia trên 5.730 sinh viên" '
       'của Nguyen, L.X. và cộng sự, đăng trên tạp chí Medicine (Q1, IF ≈ 1,5) năm 2023, DOI 10.1097/MD.0000000000033559. '
       'Khách thể nghiên cứu là 5.730 sinh viên đại học từ nhiều tỉnh thành Việt Nam, được khảo sát trực tuyến trong giai '
       'đoạn cao điểm đại dịch COVID-19. Công cụ chính là thang GAD-7 (Generalized Anxiety Disorder 7-item, Spitzer 2006) '
       'phiên bản tiếng Việt, ngưỡng cắt ≥ 10 cho lo âu mức độ vừa trở lên.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Thiết kế cắt ngang trực tuyến trên 5.730 sinh viên. Sử dụng GAD-7 sàng lọc lo âu cộng với bảng hỏi nhân khẩu xã hội: '
       'tuổi, giới tính, năm học, ngành học, nơi ở, kinh tế gia đình, tiền sử bệnh, kinh nghiệm tiếp xúc COVID-19. Phân '
       'tích hồi quy logistic đa biến để xác định các yếu tố liên quan có ý nghĩa thống kê.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ lo âu (GAD-7 ≥ 10) là 16,2 % trên toàn bộ mẫu — thấp hơn nhiều so với các nghiên cứu sàng lọc khác trong '
       'cùng giai đoạn COVID-19 ở Việt Nam (Hoàng Trung Học và cộng sự 2025 báo cáo 41,5 % bằng DASS-21). Các yếu tố nguy '
       'cơ chính được xác định là: ở ký túc xá (PR = 1,71), năm học cuối (gần thi tốt nghiệp), tiếp xúc trực tiếp với F0/F1, '
       'kinh tế gia đình giảm sút trong dịch, và thiếu hỗ trợ tâm lý từ trường.')

table(doc, ['Yếu tố', 'PR/OR', 'p'],
      [['Ở ký túc xá vs nhà', 'PR = 1,71', '< 0,001'],
       ['Năm học cuối', '↑ nguy cơ', '< 0,01'],
       ['Tiếp xúc F0/F1', '↑ nguy cơ', '< 0,001'],
       ['Kinh tế gia đình giảm', '↑ nguy cơ', '< 0,01']],
      widths=[8.0, 4.0, 3.0])

H(doc, 'Phản biện', level=2)
P(doc, 'Cỡ mẫu rất lớn n = 5.730 — thuộc nhóm khảo sát SKTT lớn nhất tại Việt Nam thời COVID. Tạp chí Medicine Q1. Tuy '
       'nhiên, mẫu thuận tiện trực tuyến → có thiên lệch tự chọn (chỉ những sinh viên có điều kiện online tham gia). Đối '
       'tượng SINH VIÊN ĐẠI HỌC, không phải VTN học sinh THCS/THPT. Tỷ lệ 16,2 % thấp hơn các nghiên cứu khác có thể do '
       'cut-off GAD-7 ≥ 10 nghiêm ngặt hơn. Cắt ngang — không đánh giá được xu hướng theo thời gian.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần nghiên cứu dọc theo dõi sức khoẻ tâm thần sinh viên trước-trong-sau COVID. So sánh GAD-7 với DASS-21 và DISC-5 '
       'trên cùng mẫu để chuẩn hoá công cụ. Phát triển mô hình hỗ trợ tâm lý cho sinh viên ký túc xá — nhóm nguy cơ cao.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Mẫu lớn, tạp chí Q1, đối tượng quan trọng, có ý nghĩa cho chính sách.', bold=True)
save(doc, 'VN23_NguyenLX_2023_Medicine.docx')

# ============================================================
# BÀI 40 — Nguyễn Thanh Truyền 2024 (Vĩnh Long n=919)
# ============================================================
print('[2/20] Bài 40: Vĩnh Long 2024')
doc = make_doc()
H(doc, 'Tóm tắt bài VN24 — Sức khoẻ tâm thần học sinh THPT tại Vĩnh Long', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Trầm cảm và các yếu tố liên quan ở học sinh trung học phổ thông tại tỉnh Vĩnh Long" của Nguyễn Thanh '
       'Truyền và cộng sự (2024), đăng trên Tạp chí Y học Việt Nam. Mẫu nghiên cứu là 919 học sinh THPT tại Vĩnh Long, '
       'sử dụng thang CES-D (Center for Epidemiologic Studies Depression) và ESSA (Educational Stress Scale for Adolescents).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Thiết kế cắt ngang trên 919 học sinh THPT tại Vĩnh Long. Công cụ: CES-D đo trầm cảm + ESSA đo căng thẳng học tập. '
       'Phỏng vấn trực tiếp tại trường. Phân tích mô tả + chi-square + hồi quy logistic đa biến.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ trầm cảm (CES-D ≥ 16) là 12,2 % — tương đối thấp so với các nghiên cứu khác tại Việt Nam. Các yếu tố nguy '
       'cơ liên quan có ý nghĩa: nghiện rượu, áp lực học tập cao theo thang ESSA, gia đình mâu thuẫn, sống xa cha mẹ. '
       'Nữ có tỷ lệ cao hơn nam.')

H(doc, 'Phản biện', level=2)
P(doc, 'Mẫu khá lớn n = 919, nhưng chỉ tại 1 tỉnh ĐBSCL — không đại diện toàn Việt Nam. Sử dụng CES-D đo TRẦM CẢM (không '
       'đo lo âu trực tiếp) — bài này phù hợp đối chiếu hơn là can thiệp trực tiếp lo âu. Tỷ lệ 12,2 % thấp có thể do '
       'cut-off CES-D ≥ 16 (nghiêm ngặt) hoặc do mẫu Vĩnh Long ít chịu áp lực thi đại học hơn các tỉnh đô thị. Không có '
       'so sánh trước-sau, không có nhóm chứng.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần bổ sung thang đo lo âu (GAD-7 hoặc DASS-21) để có bức tranh toàn diện. So sánh đa vùng (đô thị vs nông thôn). '
       'Nghiên cứu dọc theo nhóm có nguy cơ cao (nghiện rượu, gia đình mâu thuẫn).')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐ Trung bình–Khá. Mẫu khá, tạp chí trong nước, hữu ích cho dữ liệu vùng ĐBSCL.', bold=True)
save(doc, 'VN24_VinhLong_2024.docx')

# ============================================================
# BÀI 41 — Praptomojati & Hartanto 2024 (CA-CBT ĐNA SR)
# ============================================================
print('[3/20] Bài 41: Praptomojati 2024 CA-CBT SR')
doc = make_doc()
H(doc, 'Tóm tắt bài QT37 — Tổng quan hệ thống về CBT thích ứng văn hoá cho lo âu tại Đông Nam Á', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"A systematic review of Culturally Adapted Cognitive Behavioral Therapy (CA-CBT) for anxiety disorders in Southeast '
       'Asia" của Praptomojati, A. & Hartanto, A. (2024), đăng trên Asian Journal of Psychiatry vol. 92, tr. 103896 (Q1, '
       'IF ≈ 12 — tạp chí tâm thần châu Á uy tín nhất). DOI 10.1016/j.ajp.2023.103896. Phạm vi: 11 nước Đông Nam Á (Indonesia, '
       'Malaysia, Philippines, Thái Lan, Singapore, Việt Nam, ...). Tổng quan tài liệu hệ thống các nghiên cứu CBT đã được '
       'thích ứng văn hoá cho điều trị rối loạn lo âu.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống trên 6 cơ sở dữ liệu: PubMed, PsycINFO, Embase, CENTRAL, GARUDA và Google Scholar. Tiêu chí '
       'lựa chọn: nghiên cứu CBT có thích ứng văn hoá (CA-CBT) cho rối loạn lo âu tại 11 nước Đông Nam Á. Phân loại theo '
       'khung CTAF (Cultural Treatment Adaptation Framework): thích ứng ngoại vi (peripheral — ngôn ngữ, hình ảnh) so với '
       'thích ứng cốt lõi (deep — niềm tin, hệ giá trị, tôn giáo).')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, '7 nghiên cứu được chọn cuối cùng: 1 RCT, 3 nghiên cứu bán thực nghiệm và 3 ca lâm sàng. 2 nghiên cứu thực hiện '
       'thích ứng đa thành phần. 2 nghiên cứu sửa đổi thành phần CỐT LÕI bằng tích hợp tôn giáo (Hồi giáo Indonesia, Phật '
       'giáo Thái Lan) hoặc ngôn ngữ địa phương. PHÁT HIỆN QUAN TRỌNG: trong toàn bộ tổng quan, KHÔNG có nghiên cứu nào '
       'từ Việt Nam — Việt Nam là một trong những nước có khoảng trống nghiên cứu CA-CBT trong khu vực.')

H(doc, 'Phản biện', level=2)
P(doc, 'Asian Journal of Psychiatry Q1 IF ≈ 12 — tạp chí tâm thần châu Á uy tín nhất. Tổng quan hệ thống DUY NHẤT về '
       'CA-CBT lo âu tại Đông Nam Á — rất phù hợp đề tài Việt Nam. Khung CTAF giúp phân loại rõ thích ứng ngoại vi vs cốt '
       'lõi. Hạn chế: chỉ 7 nghiên cứu được tìm thấy → bằng chứng còn ít. Đa số là bán thực nghiệm hoặc ca lâm sàng, chỉ '
       '1 RCT — chất lượng bằng chứng vừa phải. Không có dữ liệu meta-analysis định lượng.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần phát triển CA-CBT phiên bản tiếng Việt với thích ứng văn hoá Á Đông: hệ giá trị Khổng giáo, gia đình mở rộng, '
       'tôn giáo (Phật giáo, Công giáo). Đặc biệt phù hợp với phát hiện của Dong, Wang & Lin 2025 (PLOS ONE) về "fear of '
       'letting down others" ở 60,3 % học sinh Trung Quốc — một cấu trúc tâm lý đặc thù Á Đông cần can thiệp riêng. Cần RCT '
       'chất lượng cao tại Việt Nam để lấp khoảng trống.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SR Q1, định hướng quan trọng cho dự án.', bold=True)
save(doc, 'QT37_Praptomojati_CA-CBT_SEA_2024.docx')

# ============================================================
# BÀI 42 — De Silva 2024 (Sri Lanka cluster RCT)
# ============================================================
print('[4/20] Bài 42: De Silva 2024 Sri Lanka cluster RCT')
doc = make_doc()
H(doc, 'Tóm tắt bài QT38 — Cluster RCT can thiệp CBT tại trường ở Sri Lanka', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Effectiveness of a cognitive behavioural therapy (CBT)-based intervention for reducing anxiety among adolescents '
       'in the Colombo District, Sri Lanka: A cluster randomised controlled trial" của De Silva, S. và cộng sự (2024), '
       'đăng trên Child and Adolescent Psychiatry and Mental Health vol. 18, article 108, tr. 1–12 (Q1, IF ≈ 4,0). DOI '
       '10.1186/s13034-024-00799-9. Open Access. Mẫu: 720 học sinh lớp 9 (~ 14 tuổi) tại 36 trường THCS khu vực Colombo, '
       'Sri Lanka. Đây là một trong số ít cluster RCT về can thiệp CBT trường tại các nước LMIC Nam Á.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Thiết kế cluster RCT — phân ngẫu nhiên TRƯỜNG học (không phải học sinh): 18 trường can thiệp + 18 trường đối chứng. '
       'Mỗi nhánh 360 học sinh, tổng cộng 720 học sinh lớp 9. Can thiệp gồm 8 phiên CBT mỗi tuần, 40 phút/phiên, do '
       'GIÁO VIÊN cung cấp sau khi được đào tạo. Chương trình bao gồm tâm lý giáo dục, kỹ thuật thở, tái cấu trúc nhận thức, '
       'kỹ năng giải quyết vấn đề và kỹ năng đối phó. Theo dõi 3 tháng sau can thiệp.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ mất mẫu < 1 % — tuân thủ rất cao, chứng tỏ tính khả thi của mô hình. Lo âu giảm có ý nghĩa thống kê sau 3 '
       'tháng: β = −0,096 (KTC 95 %: −0,186 đến −0,005, p = 0,038). Tự trọng (self-esteem) tăng có ý nghĩa: β = +0,811 '
       'sau can thiệp.')

table(doc, ['Kết quả', 'Beta', 'KTC 95 %', 'p'],
      [['Lo âu giảm', '−0,096', '−0,186 đến −0,005', '0,038'],
       ['Tự trọng tăng', '+0,811', '—', '< 0,01'],
       ['Mất mẫu', '< 1 %', '—', '—']],
      widths=[5.0, 3.0, 4.5, 2.5])

H(doc, 'Phản biện', level=2)
P(doc, 'Sri Lanka là quốc gia LMIC Nam Á — bối cảnh nguồn lực hạn chế tương tự Việt Nam. Cluster RCT là thiết kế MẠNH '
       'NHẤT cho can thiệp trường học. Mô hình "GV cung cấp CBT" rẻ tiền, có thể nhân rộng — đây là MÔ HÌNH KHẢ THI nhất '
       'cho VN trong số tất cả các nghiên cứu được tổng hợp. Hạn chế: chỉ 1 quận (Colombo District), beta lo âu nhỏ (−0,096) '
       'mặc dù p < 0,05 — hiệu lực lâm sàng cần đánh giá thêm. Theo dõi chỉ 3 tháng — chưa rõ duy trì lâu dài.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Cần lặp lại tại Việt Nam với cluster RCT tương tự — đào tạo giáo viên cố vấn tâm lý hoặc giáo viên môn tâm lý. '
       'Theo dõi 6, 12 tháng để đánh giá duy trì. Mở rộng đa vùng (đô thị, nông thôn, dân tộc thiểu số). So sánh GV vs '
       'chuyên gia tâm lý cung cấp can thiệp.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao. Cluster RCT 720 HS, Q1, mô hình khả thi cho LMIC.', bold=True)
save(doc, 'QT38_DeSilva_SriLanka_RCT_2024.docx')

# ============================================================
# BÀI 43 — Xian 2024 NMA SAD
# ============================================================
print('[5/20] Bài 43: Xian 2024 NMA SAD')
doc = make_doc()
H(doc, 'Tóm tắt bài QT39 — Network Meta-Analysis can thiệp tâm lý cho lo âu xã hội', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Psychological interventions for social anxiety disorder in children and adolescents: A systematic review and '
       'network meta-analysis" của Xian, J., Zhang, Y. & Jiang, B. (2024), đăng trên Journal of Affective Disorders vol. '
       '365, tr. 614–627 (Q1, IF ≈ 6,6). DOI 10.1016/j.jad.2024.08.097. PROSPERO CRD42023476829. Phân tích tổng hợp 30 RCT '
       'với 1.547 trẻ em và vị thành niên có rối loạn lo âu xã hội (SAD), so sánh 9 liệu pháp tâm lý + 3 nhóm đối chứng.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống PubMed, Embase, Cochrane, Web of Science. Chỉ chọn RCT. NMA Bayesian sử dụng R Studio. Đánh giá '
       'chất lượng bằng GRADE. So sánh xếp hạng các can thiệp bằng chỉ số SUCRA (Surface Under the Cumulative Ranking Curve). '
       'Tổng cộng 30 RCT được đưa vào phân tích cuối cùng.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'iCBT (CBT qua internet) đứng HẠNG 1 với SUCRA = 71,2 % — bằng chứng mạnh nhất cho can thiệp SAD ở VTN. gCBT (CBT '
       'nhóm) đứng hạng 2 với SUCRA = 68,4 %. Hai phát hiện phụ quan trọng: (1) iCBT cho hiệu quả tốt nhất với TRẦM CẢM '
       'kèm theo (SUCRA 92,2 % — "hai trong một"); (2) gCBT xếp hạng 1 cho cải thiện CHỨC NĂNG xã hội (SUCRA 89,6 %). '
       'Như vậy, kết hợp iCBT (giảm triệu chứng) với gCBT (phục hồi chức năng) có thể là chiến lược tối ưu.')

table(doc, ['Hạng', 'Liệu pháp', 'SUCRA (%)', 'Đặc điểm'],
      [['1', 'iCBT (internet CBT)', '71,2', 'Tự học online + hướng dẫn'],
       ['2', 'gCBT (CBT nhóm)', '68,4', 'Nhóm 6–12 HS, 8–16 tuần — KHẢ THI nhất'],
       ['3', 'I-CBT (CBT cá nhân)', '66,0', 'Tốn chuyên gia — khó cho LMIC'],
       ['4', 'SET-C', '~ 60', 'Kỹ năng XH + phơi nhiễm'],
       ['5', 'IPT', '~ 55', 'Liệu pháp giữa cá nhân'],
       ['—', 'WL/NT/PBO', '< 20', 'Đối chứng']],
      widths=[1.2, 5.0, 2.5, 6.5])

H(doc, 'Phản biện', level=2)
P(doc, 'NMA cho phép so sánh GIÁN TIẾP tất cả can thiệp — bằng chứng MẠNH NHẤT về xếp hạng có thể thu được không cần thử '
       'nghiệm head-to-head. iCBT hạng 1 cho SAD — nên ưu tiên phát triển app tiếng Việt. gCBT hạng 2 nhưng có ưu thế đặc '
       'biệt cho phục hồi chức năng — phù hợp triển khai ở trường. Hạn chế: chỉ tập trung SAD (không phải GAD); 30 RCT '
       'chủ yếu từ phương Tây và Đông Á phát triển (Trung Quốc, Hàn Quốc, Nhật Bản); 0 RCT từ các nước LMIC ASEAN.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Phát triển app iCBT tiếng Việt cho SAD — phù hợp xếp hạng số 1. Thử nghiệm gCBT nhóm tại trường THCS/THPT Việt '
       'Nam. RCT đối đầu giữa các can thiệp top 3 trong bối cảnh LMIC. Mở rộng NMA cho GAD và các rối loạn lo âu khác.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao. NMA 30 RCT Q1, định hướng can thiệp SAD cho VTN.', bold=True)
save(doc, 'QT39_Xian_NMA_SAD_2024.docx')

print('\n5/20 done. Continuing with bai 44-58...')

# Tiếp tục với 15 bài còn lại trong file riêng để giữ độ dài quản lý được
