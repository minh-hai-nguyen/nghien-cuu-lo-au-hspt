# -*- coding: utf-8 -*-
"""Tao tom tat VN030 Happy House RCT — bo sung khoang trong QA tim ra"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)

# ============================================================
# VN030 — Tran et al. 2023 — Happy House RCT
# ============================================================
doc = make_doc()
H(doc, 'Tóm tắt bài VN030 — Happy House: RCT can thiệp SKTT trường học cho HS lớp 10 Hà Nội', level=1)
P(doc, '⚠ LƯU Ý: Bài này chưa có bản dịch đầy đủ trong 03_Ban-dich/. Bản tóm tắt dựa trên '
       'abstract + bài đã đọc qua web.', italic=True)

P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"School-based universal mental health promotion intervention for adolescents in '
       'Vietnam: Two-arm, parallel, controlled trial" của Tran, T.D., Nguyen, H., Shochet, I., '
       'Nguyen, N., La, N., Wurfl, A., Orr, J., Nguyen, H., Stocker, R. & Fisher, J. (2023), '
       'đăng trên Cambridge Prisms: Global Mental Health vol. 10, e58 (10/2023). Link PMC: '
       'https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/')
P(doc, 'Đơn vị hợp tác Việt – Úc: Monash University (Úc), Queensland University of Technology '
       '(Úc) + đối tác Hà Nội. Khách thể: 1.084 học sinh lớp 10 (531 nhóm can thiệp + 552 '
       'nhóm đối chứng) tại 8 trường THPT Hà Nội (4 can thiệp + 4 đối chứng).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Cluster RCT hai nhánh song song có đối chứng. Phân ngẫu nhiên cấp TRƯỜNG (8 trường → '
       '4 + 4). Can thiệp: chương trình "Happy House" — phiên bản thích ứng văn hoá Việt Nam '
       'của RAP-A (Resourceful Adolescent Program), 6 buổi mỗi buổi 90 phút trong 6 tuần liền. '
       'Cơ chế: kết hợp CBT + IPT (interpersonal psychotherapy). Nội dung: positive self-talk, '
       'giữ bình tĩnh, giải quyết vấn đề, mạng lưới hỗ trợ. Người cung cấp: GIÁO VIÊN đã được '
       'đào tạo chuyên môn. Mỗi buổi gồm 2 tiết học liên tiếp.')
P(doc, 'Đo lường: trầm cảm (chính), phúc lợi tâm lý, coping self-efficacy. Thời điểm: trước '
       'can thiệp, 2 tuần sau, 6 tháng sau. Khảo sát bằng giấy + online. Phân tích logistic '
       'regression và linear mixed-effects models.')

H(doc, 'Kết quả nghiên cứu', level=2)
table(doc,
    ['Kết quả', '2 tuần sau (post)', '6 tháng sau (follow-up)'],
    [
        ['Trầm cảm (Cohen d)', 'd = 0,11 (p = 0,011) *', 'KHÔNG còn ý nghĩa'],
        ['Tỷ lệ trầm cảm lâm sàng', 'OR = 0,56 (p = 0,027) *', 'KHÔNG còn ý nghĩa'],
        ['Phúc lợi tâm lý', 'd = 0,13 (p < 0,05) *', 'KHÔNG còn ý nghĩa'],
        ['Coping self-efficacy', 'd = 0,17 – 0,26 *', 'd = 0,17 – 0,26 * DUY TRÌ'],
    ],
    widths=[5.5, 5.0, 5.0])
P(doc, '* có ý nghĩa thống kê', italic=True, size=10)

P(doc, 'Bảng tổng kết:', bold=True)
P(doc, '• Hiệu ứng trên trầm cảm NHỎ (Cohen d = 0,11) — phù hợp với cảnh báo của Zhang, '
       'Liang & Kang 2026 (Bayesian MA 31 RCT) rằng universal CBT cho hiệu quả nhỏ.')
P(doc, '• Hiệu ứng "fade-out" ở 6 tháng — cả trầm cảm, phúc lợi đều không còn ý nghĩa thống kê. '
       'Có thể do COVID-19 gián đoạn HOẶC do liều can thiệp quá ngắn (cần booster).')
P(doc, '• Coping self-efficacy DUY TRÌ ở 6 tháng — đây là phát hiện tích cực: kỹ năng ứng phó '
       'có thể được học và giữ lại lâu dài, dù triệu chứng quay lại.')

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh:', bold=True)
P(doc, '• ĐÂY LÀ RCT CLUSTER ĐẦU TIÊN tại Việt Nam về can thiệp SKTT học đường cho VTN — lấp '
       'khoảng trống nghiên cứu mà các tổng quan trước (Praptomojati 2024 SR ĐNA: 0/7 NC từ VN; '
       'Menon 2025 scoping LMIC ĐÁ-TBD) đã chỉ ra.')
P(doc, '• Cỡ mẫu rất lớn (n = 1.084) — đủ power cho phân tích đa biến và subgroup.')
P(doc, '• RAP-A có cơ sở bằng chứng quốc tế mạnh từ Úc + được thích ứng văn hoá Việt Nam '
       'một cách hệ thống (xem Nguyen et al. 2022 PLOS ONE 17(8):e0271959 về quá trình adaptation).')
P(doc, '• Mô hình hợp tác Việt – Úc (Monash, QUT) — bài học cho hợp tác quốc tế.')
P(doc, '• Giáo viên cung cấp — phù hợp bối cảnh LMIC thiếu chuyên gia tâm lý.')

P(doc, 'Hạn chế:', bold=True)
P(doc, '• Đây là can thiệp UNIVERSAL (toàn lớp), không phải TARGETED cho HS có triệu chứng. '
       'Hiệu ứng nhỏ d = 0,11 có thể do dilution.')
P(doc, '• Hiệu ứng KHÔNG CÒN ý nghĩa ở 6 tháng — gợi ý cần BOOSTER SESSIONS.')
P(doc, '• Đo TRẦM CẢM chính, KHÔNG đo LO ÂU riêng — báo cáo VTN VN cho lo âu vẫn chưa có '
       'RCT chuyên biệt.')
P(doc, '• Chỉ Hà Nội + lớp 10 (không phải lớp 11-12 và THCS) — chưa đại diện toàn VN.')
P(doc, '• KHÔNG có nhóm active control — chỉ là usual curriculum, khó loại trừ hiệu ứng chú ý.')
P(doc, '• Có thể bị ảnh hưởng bởi COVID-19 (dữ liệu thu thập trong giai đoạn restrictions).')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, '(1) Lặp lại với THIẾT KẾ TARGETED — chỉ chọn HS có GAD-7 ≥ 8 hoặc DASS-Y lo âu ≥ 8 '
       '— có thể có hiệu ứng lớn hơn.')
P(doc, '(2) Thêm BOOSTER SESSIONS ở 3 và 6 tháng để duy trì hiệu quả.')
P(doc, '(3) Đo CẢ trầm cảm + LO ÂU + coping self-efficacy + resilience.')
P(doc, '(4) Mở rộng địa bàn ra TPHCM, Đà Nẵng, miền Trung nông thôn, miền núi DTTS.')
P(doc, '(5) Mở rộng tuổi: lớp 11, 12 (gần Gao Kao) và THCS (lớp 6-9).')
P(doc, '(6) So sánh GIÁO VIÊN cung cấp vs CHUYÊN GIA TÂM LÝ vs Y TÁ HỌC ĐƯỜNG.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Cluster RCT đầu tiên tại VN cho VTN, mẫu lớn, '
       'tác giả Monash–QUT uy tín. Hạn chế: universal interventions với hiệu ứng nhỏ, không đo '
       'lo âu riêng. PHẢI ĐƯỢC TRÍCH DẪN trong mọi báo cáo can thiệp VN từ nay trở đi.',
  bold=True)

out_path = os.path.join(TT_DIR, 'VN030_Tran_HappyHouse_RCT_GMH_2023.docx')
doc.save(out_path)
d = Document(out_path)
chars = sum(len(p.text) for p in d.paragraphs)
print(f'Saved: {os.path.basename(out_path)}')
print(f'Chars: {chars}, Tables: {len(d.tables)}')
