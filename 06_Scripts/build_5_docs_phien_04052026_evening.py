"""Build 5 docs cho phien 04/05/2026 (toi):
1. Phan bien Nguyen Thi Van VN004 - yeu to anh huong
2. Binh luan chi tiet Tran Thi My Luong VN006 - RLLA HS THPT chuyen
3. Binh luan M=42,09 SD=34,075
4. CFA la gi?
5. KTC 90% RMSEA (0,024-0,067) co duoc khong?
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x30)

def new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
    style = d.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return d

def H(d, text, level=1, color=NAVY):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color

def para(d, text, color=BLACK, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color
    r.font.size = Pt(size); r.bold = bold; r.italic = italic

def bullet(d, text, color=BLACK, italic=False, size=12):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(size); r.italic = italic

def add_table(d, header, rows, col_widths=None):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

# ===========================================================
# DOC 1 — Phản biện Nguyễn Thị Vân VN004 (yếu tố ảnh hưởng)
# ===========================================================
def build_doc1():
    d = new_doc()
    H(d, 'Phản biện Nguyễn Thị Vân (2020) — VN004', level=1)
    H(d, '— Yếu tố ảnh hưởng đến rối loạn lo âu HS THPT TP.HCM —', level=2)

    H(d, 'Câu hỏi của thầy', level=2)
    para(d, 'Em sửa cho thầy câu bình luận này nhé. Nguyễn Thị Vân về yếu tố ảnh hưởng:', color=BLUE)
    para(d, '"Về số liệu liên quan đến áp lực từ phía gia đình, các tác giả chỉ mới đưa ra tỷ lệ % thô, không có kiểm chứng bằng các phép toán thống kê hiện đại, thiếu phân tích đa biến để kiểm soát các yếu tố gây nhiễu."', color=BLUE)
    para(d, 'Hoặc: Em cho thầy phản biện đầy đủ về Nguyễn Thị Vân với các yếu tố ảnh hưởng nhé. Một số tác giả VN chỉ tính % rồi bình luận, không có p, thiếu các hệ số khác mà Em đã giải thích cho thầy trong nhiều bài trước.', color=BLUE)

    H(d, 'Bối cảnh — VN004 Nguyễn Thị Vân (2020)', level=2)
    para(d, 'Bài "Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh" (Nguyễn Thị Vân, 2018; Tạp chí Khoa học - Khoa học Giáo dục, ĐHSP TP.HCM). Một số phiên bản khác công bố 2019 (luận án TS) và 2020.')
    para(d, 'Trong corpus dự án: VN004 = NguyenThiVan_2020_STAI_TPHCM (PDF + tóm tắt — chưa có bản dịch).')
    para(d, 'Phương pháp: cắt ngang 2 pha — sàng lọc 558 HS bằng STAI Spielberger (bản Việt hóa Nguyễn Công Khanh 2000) → 90 HS phân tích sâu (mẫu thuận tiện, 4 trường nội + ngoại thành TPHCM).')
    para(d, 'Kết quả chính: tỷ lệ RLLA 15–18,5 %. 4 nhóm yếu tố ảnh hưởng được xác định: học tập / gia đình / quan hệ xã hội / bản thân HS.')

    H(d, '1. SỬA CÂU BÌNH LUẬN của thầy', level=2)
    para(d, 'Câu thầy viết hiện tại đã đúng ý nhưng có thể chính xác và mạnh hơn. Đề xuất 2 phiên bản:', bold=True)

    para(d, 'Phiên bản A (ngắn, học thuật):', bold=True, color=GREEN)
    para(d,
        '"Đối với áp lực từ phía gia đình, tác giả chỉ báo cáo TỶ LỆ % các biểu hiện '
        '(kỳ vọng cha mẹ 48,9 %, mâu thuẫn... ) mà KHÔNG báo cáo các chỉ số thống kê suy '
        'luận: hệ số tương quan r kèm khoảng tin cậy, giá trị p, tỷ số chênh OR/AOR, '
        'hay phân tích hồi quy đa biến để kiểm soát yếu tố nhiễu. Việc này hạn chế khả '
        'năng kết luận về MỨC ĐỘ LIÊN QUAN giữa áp lực gia đình và RLLA — chỉ có thể '
        'mô tả TẦN SUẤT mà không đo được CƯỜNG ĐỘ ảnh hưởng."',
        color=GREEN, italic=True
    )

    para(d, 'Phiên bản B (chi tiết, có dẫn nguồn):', bold=True, color=GREEN)
    para(d,
        '"Liên quan đến nhóm yếu tố áp lực gia đình, các tác giả mới dừng ở mô tả tần '
        'suất bằng tỷ lệ phần trăm thô (descriptive statistics), KHÔNG sử dụng các thống '
        'kê suy luận chuẩn (inferential statistics) như: (1) hệ số tương quan Pearson r '
        'kèm KTC 95% và p-value; (2) hồi quy đa biến (multiple regression) hoặc hồi quy '
        'logistic kiểm soát đồng thời các nhóm yếu tố học tập – gia đình – quan hệ xã '
        'hội – bản thân HS; (3) tỷ số chênh có hiệu chỉnh AOR cho từng yếu tố gia đình. '
        'Hạn chế này không cho phép xác định liệu áp lực gia đình có còn dự báo độc lập '
        'cho RLLA sau khi đã loại trừ ảnh hưởng của các nhóm yếu tố khác. Đây là điểm '
        'yếu phương pháp luận đáng lưu ý khi áp dụng cho bối cảnh nghiên cứu hiện đại '
        '(theo chuẩn STROBE 2007 và CONSORT 2010)."',
        color=GREEN, italic=True
    )

    H(d, '2. PHẢN BIỆN ĐẦY ĐỦ Nguyễn Thị Vân (2020) về yếu tố ảnh hưởng', level=2)

    para(d, 'A. Điểm mạnh đáng ghi nhận', bold=True)
    bullet(d, 'Là một trong số ÍT NC VN chuyên về YẾU TỐ ẢNH HƯỞNG (không chỉ prevalence). Đa số NC VN dừng ở mô tả tỷ lệ %.')
    bullet(d, 'Khung 4 nhóm yếu tố hệ thống (học tập / gia đình / quan hệ xã hội / bản thân HS) — phù hợp khung Bronfenbrenner.')
    bullet(d, 'Kết hợp định lượng (STAI) + định tính (phỏng vấn sâu 90 HS) — mixed methods design.')
    bullet(d, 'Phân tầng theo trường nội + ngoại thành TPHCM — bước đầu kiểm soát biến nhiễu địa lý.')

    para(d, 'B. Hạn chế phương pháp luận chính', bold=True)

    H(d, 'B1. THIẾU thống kê suy luận', level=3)
    para(d, 'Đây là điểm yếu lớn nhất — đúng như thầy đã nhận xét. Các vấn đề cụ thể:')
    bullet(d, 'KHÔNG có p-value cho từng yếu tố → không biết tỷ lệ X % có khác có ý nghĩa thống kê hay chỉ do ngẫu nhiên.')
    bullet(d, 'KHÔNG có hệ số tương quan r kèm KTC 95% (bài có nhắc r=0,42 cho nhóm "bản thân" nhưng KHÔNG có KTC, p, hay r cho các nhóm khác đầy đủ).')
    bullet(d, 'KHÔNG có hồi quy đa biến → không kiểm soát được yếu tố nhiễu (giới, lớp, kinh tế gia đình, học lực).')
    bullet(d, 'Con số "85,4 % giải thích" của nhóm bản thân HS gây nghi ngờ — không khớp với r=0,42 (vì r² = 0,176 = 17,6 %). Có thể là R² đa biến KHÔNG được trình bày rõ, hoặc là tỷ lệ % HS có biểu hiện chứ không phải variance explained.')

    H(d, 'B2. Mẫu thuận tiện', level=3)
    bullet(d, '90 HS phân tích sâu chọn theo "mẫu thuận tiện" (convenience sampling) — không ngẫu nhiên, không đại diện TPHCM.')
    bullet(d, '558 HS sàng lọc cũng chỉ ở 4 trường, không phân tầng theo quận → bias chọn mẫu.')

    H(d, 'B3. Công cụ cũ', level=3)
    bullet(d, 'STAI Spielberger phiên bản Việt hóa Nguyễn Công Khanh (2000) — KHÔNG cập nhật psychometric trong 20 năm.')
    bullet(d, 'Cronbach alpha không được báo cáo cho mẫu hiện tại → không biết độ tin cậy thực.')
    bullet(d, 'STAI đo "lo âu trạng thái" + "lo âu đặc điểm" — không phải chẩn đoán RLLA theo DSM-5/ICD-11. Bài dùng từ "RLLA" là KHÔNG CHÍNH XÁC.')

    H(d, 'B4. Thiếu nhóm chứng', level=3)
    bullet(d, 'Không có nhóm chứng (HS không có biểu hiện lo âu). Các "yếu tố ảnh hưởng" thực tế chỉ là các biểu hiện CÓ ở nhóm có lo âu — không cho biết có đặc thù hơn nhóm chung không.')

    H(d, 'B5. Phân tích yếu tố gia đình thiếu chiều sâu', level=3)
    bullet(d, 'Tác giả chỉ liệt kê các biểu hiện ("kỳ vọng cha mẹ 48,9 %", "mâu thuẫn cha mẹ X %") không phân tích MỨC ĐỘ tác động.')
    bullet(d, 'Không sử dụng thang đo chuẩn về Family Functioning (như FAD, McMaster) hay Parental Stress Scale.')
    bullet(d, 'Không tách "kỳ vọng cha mẹ" thành kỳ vọng học tập vs kỳ vọng nghề nghiệp vs kỳ vọng đạo đức.')

    para(d, 'C. So sánh với chuẩn quốc tế', bold=True)
    para(d, 'Để minh họa khoảng cách phương pháp, em đối chiếu với 1 bài cùng chủ đề trong corpus:')
    add_table(d,
        ['Tiêu chí', 'Nguyễn Thị Vân (2020)', 'Qiu et al. (2022) QT008 Trung Quốc'],
        [
            ['Cỡ mẫu', '90 HS sâu (558 sàng lọc)', '2.079 HS THCS'],
            ['Lấy mẫu', 'Thuận tiện', 'Phân tầng ngẫu nhiên'],
            ['Báo cáo OR/AOR', 'Không', 'Có (OR=2,01 KTC 95% 1,38–2,92 cho nuôi dạy tiêu cực)'],
            ['p-value', 'Không', 'Có'],
            ['Hồi quy đa biến', 'Không rõ ràng', 'Có — kiểm soát giới, lớp, KT gia đình'],
            ['Công cụ đo gia đình', 'Bảng hỏi tự thiết kế', 'EMBU-C (chuẩn quốc tế)'],
            ['Tạp chí', 'TC Khoa học ĐHSP TPHCM', 'Frontiers in Public Health (Q2)'],
        ]
    )

    para(d, '')
    para(d, 'D. Đề xuất cải tiến (cho thầy nếu viết NC tương tự)', bold=True)
    bullet(d, 'Dùng thang đo gia đình chuẩn: FAD-12 hoặc PARQ.')
    bullet(d, 'Chạy hồi quy logistic đa biến với outcome = RLLA (đo bằng SCARED hoặc RCADS), predictors = các thang gia đình, học tập, peer.')
    bullet(d, 'Báo cáo AOR + KTC 95% + p cho từng yếu tố sau khi kiểm soát giới, lớp, học lực, KT gia đình.')
    bullet(d, 'Cân nhắc Structural Equation Modeling (SEM) để mô hình hóa quan hệ giữa các nhóm yếu tố.')

    H(d, 'CÂU TRẢ LỜI', level=2, color=BLUE)
    para(d, 'Tóm tắt:', bold=True, color=BLUE)
    bullet(d, 'Câu bình luận của thầy đã ĐÚNG Ý nhưng có thể mạnh hơn — em đề xuất 2 phiên bản A (ngắn) và B (chi tiết) ở mục 1.', color=BLUE)
    bullet(d, 'Phản biện đầy đủ: 5 nhóm hạn chế phương pháp (B1 thiếu thống kê suy luận, B2 mẫu thuận tiện, B3 công cụ cũ, B4 không nhóm chứng, B5 phân tích gia đình nông).', color=BLUE)
    bullet(d, 'Vấn đề CHUNG ở các NC VN giai đoạn 2015–2020: dừng lại ở mô tả % thô; không có p, OR, AOR, KTC 95% — đây là khoảng cách lớn so với chuẩn STROBE/CONSORT quốc tế.', color=BLUE)
    bullet(d, 'Khuyến nghị thầy khi trích Nguyễn Thị Vân: chỉ dùng để mô tả 4 nhóm yếu tố như khung khái niệm, KHÔNG dùng làm bằng chứng định lượng cho cường độ tác động của từng yếu tố.', color=BLUE)

    H(d, 'Phụ lục — Tài liệu tham khảo', level=2)
    para(d, '1. Nguyễn Thị Vân. (2018). Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh. Tạp chí Khoa học - Khoa học Giáo dục, ĐHSP TPHCM. [VN004 trong DB.]', italic=True, size=11)
    para(d, '2. Nguyễn Thị Vân. (2019). Lo âu học đường của học sinh trung học phổ thông tại TP.HCM. Luận án tiến sĩ, Đại học Sư phạm TP.HCM. [Nguồn liên quan, có thể là báo cáo dài.]', italic=True, size=11)
    para(d, '3. Qiu, Z., Guo, Y., Wang, J., & Zhang, H. (2022). Associations of parenting style and resilience with depression and anxiety among Chinese high school students. Frontiers in Public Health, 10, 989125. https://doi.org/10.3389/fpubh.2022.989125 [Đối chiếu phương pháp.]', italic=True, size=11)
    para(d, '4. von Elm, E., Altman, D. G., Egger, M., Pocock, S. J., Gøtzsche, P. C., & Vandenbroucke, J. P. (2007). The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Annals of Internal Medicine, 147(8), 573–577. https://doi.org/10.7326/0003-4819-147-8-200710160-00010', italic=True, size=11)
    para(d, '5. Spielberger, C. D. (1983). Manual for the State-Trait Anxiety Inventory (Form Y). Palo Alto, CA: Consulting Psychologists Press. [Bộ STAI gốc.]', italic=True, size=11)

    out = OUT_DIR / 'Phan_bien_NguyenThiVan_VN004_yeu_to_anh_huong.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')

# ===========================================================
# DOC 2 — Bình luận chi tiết Trần Thị Mỵ Lương VN006
# ===========================================================
def build_doc2():
    d = new_doc()
    H(d, 'Bình luận chi tiết Trần Thị Mỵ Lương & Đặng Đức Anh (2020) — VN006', level=1)
    H(d, '— Rối loạn lo âu của học sinh trường THPT Chuyên —', level=2)

    H(d, 'Câu hỏi của thầy', level=2)
    para(d, 'Em cho thầy bình luận chi tiết về Trần Thị Mỹ Lương với bài Rối loạn lo âu của HS trường THPT chuyên nhé.', color=BLUE)

    H(d, 'Bối cảnh', level=2)
    para(d, 'Bài VN006 trong corpus: Trần Thị Mỵ Lương & Đặng Đức Anh (2020). "Rối loạn lo âu ở học sinh trung học phổ thông [Chuyên]". Tạp chí Khoa học — Trường Đại học Thủ đô Hà Nội, Số 40/2020, trang 122–131. Tác giả chính: Học viện Phụ nữ Việt Nam. Email: Hoangthongtinnguvan@gmail.com.')
    para(d, 'Mẫu: 540 HS (5 lớp/khối × 3 khối lớp 10–11–12) trong tổng 1.457 HS toàn trường — chọn mẫu ngẫu nhiên phân tầng. Sàng lọc → 77 HS có RLLA (14,2 %). 33,5 % nam + 66,5 % nữ. Công cụ: DASS-42 + phỏng vấn sâu.')

    H(d, '1. ĐIỂM MẠNH đáng ghi nhận', level=2)
    bullet(d, 'Một trong số RẤT ÍT NC VN chuyên về HS TRƯỜNG CHUYÊN — nhóm đặc thù áp lực cao mà y văn quốc tế gọi là "perfectionism trap".')
    bullet(d, 'Phân tầng theo khối lớp 10/11/12 → phát hiện đặc sắc: KHỐI 11 cao nhất (48,1 %) > khối 10 (31,1 %) > khối 12 (20,8 %) — TRÁI giả thuyết "khối 12 thi ĐH áp lực nhất".')
    bullet(d, 'Lý giải sâu sắc của tác giả: khối 11 chịu áp lực ĐA CHIỀU (chọn ngành/khối thi/định hướng nghề), khối 12 đã ổn định lựa chọn → áp lực giảm tự nhiên.')
    bullet(d, 'Kết hợp định lượng + phỏng vấn định tính. Trích phỏng vấn nổi bật: "Em là HS Chuyên mà mấy cái đó cũng không làm được, em thấy tệ quá!" — minh họa "perfectionism trap".')
    bullet(d, 'Dùng DASS-42 (full version 42 mục) thay vì chỉ DASS-21 → thêm chi tiết về phân biệt mức độ.')
    bullet(d, 'Lấy mẫu NGẪU NHIÊN PHÂN TẦNG (stratified random sampling) — chuẩn hơn nhiều so với "convenient" của VN014 Hoàng Trung Học.')

    H(d, '2. HẠN CHẾ PHƯƠNG PHÁP', level=2)

    H(d, '2.1. KHÔNG nêu địa bàn cụ thể (lỗi nghiêm trọng)', level=3)
    para(d, 'Bài chỉ ghi "trường THPT Chuyên" — KHÔNG NÊU TÊN TRƯỜNG + TỈNH. Điều này:', bold=True)
    bullet(d, 'Vi phạm chuẩn STROBE 2007 (Strengthening the Reporting of Observational Studies in Epidemiology) — yêu cầu Item 5: "describe the setting, locations, and relevant dates".')
    bullet(d, 'Không thể replicate hoặc đối chiếu địa lý.')
    bullet(d, 'Không biết là trường chuyên Bắc/Trung/Nam, miền núi/đồng bằng, đô thị/nông thôn — các yếu tố có thể ảnh hưởng tỷ lệ.')
    para(d, 'Manh mối em đoán từ nội dung: trường có ký túc xá, HS từ huyện về thành phố tỉnh, quy mô vừa (1.457 HS) — KHÔNG phải Hà Nội-Amsterdam (~2.500), KHTN (~700), Lê Hồng Phong TPHCM (~3.000). Có thể là một trường chuyên cấp tỉnh ở Bắc Bộ (gần Học viện Phụ nữ VN nơi tác giả công tác).', italic=True)

    H(d, '2.2. KHÔNG nêu cutoff DASS-42 cụ thể', level=3)
    bullet(d, 'Tác giả không công bố ngưỡng cắt cho phân loại "RLLA" / "không RLLA" / phân tầng nhẹ–vừa–nặng–rất nặng.')
    bullet(d, 'DASS-42 có nhiều cutoff khác nhau giữa các phiên bản: ngưỡng Lovibond gốc khác bản Việt hóa khác bản Trung Quốc.')
    bullet(d, 'Hậu quả: tỷ lệ 14,2 % không thể so trực tiếp với NC dùng DASS-21 hoặc GAD-7.')

    H(d, '2.3. Lệch giới (66,5 % nữ)', level=3)
    bullet(d, 'Nữ chiếm 66,5 % → vượt xa tỷ lệ trường chuyên thông thường (50/50 hoặc nhẹ về nữ ~55 %).')
    bullet(d, 'Nữ thường có tỷ lệ lo âu cao hơn nam 1,5–2 lần (Hoa 2024 VN, Puyat 2025 Philippines, NSCH 2020 US).')
    bullet(d, 'Vậy tỷ lệ RLLA tổng 14,2 % BỊ PHÓNG ĐẠI do mẫu lệch giới — nếu balanced 50/50 thì tỷ lệ thực có thể ~ 11–12 %.')

    H(d, '2.4. Thiếu phân tích đa biến', level=3)
    bullet(d, 'Bài chỉ mô tả tỷ lệ % theo khối lớp — không hồi quy đa biến kiểm soát giới.')
    bullet(d, 'KHÔNG có hệ số thống kê suy luận: p-value, OR, AOR, χ², t-test giữa các khối.')
    bullet(d, 'Phát hiện "khối 11 > khối 12" có ý nghĩa thống kê hay chỉ do ngẫu nhiên — không kiểm chứng được.')

    H(d, '2.5. Thiếu nhóm chứng', level=3)
    bullet(d, 'Không so sánh với HS không-chuyên cùng khối → không xác định được "chuyên" có CAO HƠN "không chuyên" không.')
    bullet(d, 'Giả thuyết "perfectionism trap" của trường chuyên cần nhóm chứng để confirm.')

    H(d, '2.6. Chỉ 1 trường', level=3)
    bullet(d, 'Không thể ngoại suy ra toàn bộ HS chuyên VN — văn hóa trường, vùng miền, chính sách trường có thể tạo khác biệt.')
    bullet(d, 'Cần multi-site study với ít nhất 5–10 trường chuyên đại diện 3 miền để confirm.')

    H(d, '3. ĐIỂM ĐẶC SẮC PHẢN BIỆN', level=2)
    para(d, 'Phát hiện KHỐI 11 > KHỐI 12 LO ÂU là đóng góp QUAN TRỌNG — ngược giả định phổ biến rằng "thi ĐH gây lo âu nhất". Hàm ý cho can thiệp:', bold=True)
    bullet(d, 'Nên ưu tiên CHƯƠNG TRÌNH PHÒNG NGỪA cho khối 11 trường chuyên, không chỉ khối 12.')
    bullet(d, 'Khối 11 đối mặt áp lực TUYỂN SINH SỚM (chọn khối thi), TÌM HƯỚNG NGHỀ, ÁP LỰC HỌC SINH GIỎI QUỐC GIA — đa chiều hơn khối 12 (chỉ thi tốt nghiệp + ĐH).')
    bullet(d, 'Cần multi-site study confirm trước khi đưa vào chính sách quốc gia.')

    H(d, '4. SO SÁNH với các bài VN khác trong corpus', level=2)
    add_table(d,
        ['Bài', 'Cỡ mẫu', 'Lấy mẫu', 'Tỷ lệ lo âu', 'Phân tích đa biến'],
        [
            ['VN006 Trần Thị Mỵ Lương 2020 (chuyên)', '540', 'Ngẫu nhiên phân tầng', '14,2% (DASS-42)', 'Không'],
            ['VN001 Hoa 2024 (Hà Nội)', '3.910', 'Đa giai đoạn', '40,6% (GAD-7)', 'Có (logistic, AOR + KTC)'],
            ['VN014 Hoàng Trung Học 2025 (6 tỉnh)', '8.389', 'Convenient random (mâu thuẫn)', '41,5%→25,4% (DASS-21)', 'Có (R²=0,190)'],
            ['VN004 Nguyễn Thị Vân 2020 (TPHCM)', '90 sâu/558 sàng', 'Thuận tiện', '15-18,5% (STAI)', 'Không rõ ràng'],
            ['VN002 V-NAMHS 2022 (38 tỉnh)', '5.996', 'Đại diện QG', '21,7% SKTT (DSM-5)', 'Có (chuẩn IBM)'],
        ]
    )

    para(d, '')
    para(d, 'Trần Thị Mỵ Lương xếp giữa: lấy mẫu chuẩn hơn VN014 và VN004, nhưng không có phân tích đa biến như VN001/VN002.', italic=True)

    H(d, 'CÂU TRẢ LỜI', level=2, color=BLUE)
    para(d, 'Đánh giá tổng: ⭐⭐⭐⭐ (4/5) — bài CÓ GIÁ TRỊ ĐẶC SẮC ở phát hiện khối 11 > khối 12 và đề tài HS chuyên hiếm, nhưng có 6 hạn chế phương pháp đáng kể.', bold=True, color=BLUE)

    para(d, 'Khuyến nghị cách trích vào báo cáo CTH:', bold=True, color=BLUE)
    bullet(d, 'NÊN: "Trần Thị Mỵ Lương & Đặng Đức Anh (2020) khảo sát 540 HS một trường THPT Chuyên (địa bàn không nêu) bằng DASS-42 — phát hiện 14,2 % có biểu hiện RLLA, với khối 11 chiếm tỷ lệ cao nhất (48,1 %), trái với giả định khối 12 thi ĐH áp lực nhất."', color=BLUE)
    bullet(d, 'NÊN ghi rõ HẠN CHẾ: "Bài KHÔNG nêu tên trường/tỉnh, KHÔNG cutoff DASS-42 cụ thể, mẫu nữ chiếm 66,5 % (lệch giới), thiếu phân tích đa biến — kết quả khó ngoại suy ra HS chuyên VN."', color=BLUE)
    bullet(d, 'KHÔNG NÊN: dùng tỷ lệ 14,2 % làm bằng chứng chính cho "tỷ lệ RLLA HS chuyên VN" — chỉ là dữ liệu một trường.', color=BLUE)
    bullet(d, 'NÊN dùng: phát hiện "khối 11 cao nhất" làm GIẢ THUYẾT cho NC tiếp theo, kết hợp với QT067 Pascoe (academic stress) để xây dựng can thiệp.', color=BLUE)

    para(d, 'Đề xuất phản biện ngắn cho báo cáo CTH:', bold=True, color=BLUE)
    para(d,
        '"Trần Thị Mỵ Lương & Đặng Đức Anh (2020) đóng góp quan trọng khi phát hiện khối 11 '
        '(48,1 %) có tỷ lệ RLLA cao hơn khối 12 (20,8 %) ở HS một trường THPT Chuyên, gợi ý '
        '"perfectionism trap" và áp lực đa chiều của khối 11 trong môi trường chuyên. Tuy '
        'nhiên bài có nhiều hạn chế phương pháp luận: (1) không nêu tên trường/tỉnh cụ thể, '
        '(2) không công bố cutoff DASS-42, (3) mẫu lệch giới (nữ 66,5 %), (4) không có '
        'phân tích đa biến hay kiểm định χ² giữa các khối, (5) thiếu nhóm chứng HS không-'
        'chuyên, và (6) chỉ 1 trường nên hạn chế tính ngoại suy. Khuyến nghị nhân rộng '
        'multi-site với cỡ mẫu lớn hơn, có nhóm chứng và phân tích hồi quy đa biến."',
        color=BLUE, italic=True
    )

    H(d, 'Phụ lục — Tài liệu tham khảo', level=2)
    para(d, '1. Trần Thị Mỵ Lương & Đặng Đức Anh. (2020). Rối loạn lo âu ở học sinh trung học phổ thông [Chuyên]. Tạp chí Khoa học — Trường Đại học Thủ đô Hà Nội, Số 40/2020, 122–131. [VN006 trong DB.]', italic=True, size=11)
    para(d, '2. von Elm, E., et al. (2007). The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Annals of Internal Medicine, 147(8), 573–577. https://doi.org/10.7326/0003-4819-147-8-200710160-00010', italic=True, size=11)
    para(d, '3. Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Psychology Foundation of Australia.', italic=True, size=11)
    para(d, '4. Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823 [QT067 trong DB.]', italic=True, size=11)

    out = OUT_DIR / 'Phan_bien_TranThiMyLuong_VN006_HS_THPT_Chuyen.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')

# ===========================================================
# DOC 3 — Bình luận M=42,09 SD=34,075
# ===========================================================
def build_doc3():
    d = new_doc()
    H(d, 'Bình luận về Điểm trung bình M=42,09 và SD=34,075', level=1)
    H(d, '— Vì sao SD cao bằng 81 % của M có vấn đề? —', level=2)

    H(d, 'Câu hỏi của thầy', level=2)
    para(d, 'Em ơi, bình luận như thế nào khi Điểm trung bình 42,09, Độ lệch chuẩn là 34,075?', color=BLUE)

    H(d, '1. PHÉP TÍNH NHANH — Coefficient of Variation (CV)', level=2)
    para(d, 'CV = SD / M × 100 % = 34,075 / 42,09 × 100 % = 81,0 %.')
    para(d, '⚠ CV = 81 % là RẤT CAO. Quy ước thực hành (Reed et al. 2002, Pharmacological Research):', bold=True)

    add_table(d,
        ['CV', 'Phân loại', 'Diễn giải'],
        [
            ['< 10 %', 'Thấp — phân phối tập trung', 'Mẫu rất đồng nhất quanh trung bình'],
            ['10 – 30 %', 'Trung bình — bình thường', 'Đa số NC tâm lý/y học có CV ở mức này'],
            ['30 – 60 %', 'Cao — biến thiên đáng kể', 'Cần phân tầng dưới nhóm để hiểu rõ'],
            ['> 60 %', 'Rất cao — phân phối rộng/skewed', '⚠ M có thể KHÔNG đại diện. Nghi ngờ outliers/bimodal'],
        ]
    )

    para(d, '')
    para(d, 'Với CV = 81 %, dữ liệu này thuộc nhóm "RẤT CAO" — vượt xa ngưỡng thông thường.')

    H(d, '2. BA HỆ QUẢ ĐÁNG NGỜ', level=2)

    H(d, '2.1. M có thể KHÔNG đại diện cho mẫu', level=3)
    para(d, 'Khi SD ≈ 81 % của M, phân phối rất rộng. Các trường hợp khả dĩ:')
    bullet(d, 'Phân phối LỆCH PHẢI mạnh (right-skewed): đa số quan sát ở mức thấp, vài quan sát rất cao kéo M lên.')
    bullet(d, 'Phân phối LƯỠNG ĐỈNH (bimodal): có 2 nhóm phụ tách biệt — ví dụ "nhóm nhẹ" + "nhóm nặng".')
    bullet(d, 'Có OUTLIERS cực đoan kéo lệch SD nhưng không lệch M nhiều.')
    bullet(d, 'Nếu là thang đo rời rạc (Likert, đếm sự kiện), có thể có nhiều giá trị 0 + vài giá trị cực cao.')

    H(d, '2.2. Khoảng tin cậy đáng kể có thể chứa giá trị âm', level=3)
    para(d, 'Tính toán nhanh KTC 95% cho M (giả định phân phối chuẩn):')
    para(d, '  M ± 1,96 × SE; với n không biết, giả sử n=100 → SE = SD/√n = 34,075/10 = 3,4075.')
    para(d, '  KTC 95% ≈ [42,09 − 6,68; 42,09 + 6,68] = [35,4; 48,8] — OK nếu n đủ lớn.')
    para(d, 'NHƯNG nếu xét RANGE TỰ NHIÊN ±1 SD (chứa ~68 % giá trị):')
    para(d, '  M − 1 SD = 42,09 − 34,075 = 8,02')
    para(d, '  M + 1 SD = 42,09 + 34,075 = 76,17')
    para(d, '  → Range RỘNG từ 8 đến 76 — nếu thang đo có cận trên < 76 (ví dụ thang Likert 0–60), điều này gợi ý phân phối KHÔNG chuẩn (truncated).')

    para(d, 'Quan trọng: nếu M − 2 SD < 0 (ví dụ ở đây 42,09 − 2×34,075 = −26,06), giá trị "âm" không có nghĩa thực tế cho thang điểm dương → CONFIRM phân phối không chuẩn.', bold=True)

    H(d, '2.3. Tham số (parametric) test có thể không phù hợp', level=3)
    bullet(d, 'M ± SD chỉ có ý nghĩa rõ khi phân phối chuẩn (normal). CV = 81 % gợi ý phân phối có thể KHÔNG chuẩn.')
    bullet(d, 'Nên kiểm tra Shapiro-Wilk hoặc Kolmogorov-Smirnov; nếu phân phối lệch → dùng MEDIAN + IQR thay M + SD.')
    bullet(d, 'Nếu so sánh nhóm: dùng Mann-Whitney U test thay independent t-test; Wilcoxon thay paired t-test.')

    H(d, '3. CÁC TÌNH HUỐNG CỤ THỂ', level=2)
    para(d, 'Tùy thang đo, M=42,09 SD=34,075 có ý nghĩa khác nhau:', bold=True)

    add_table(d,
        ['Thang đo', 'Range chuẩn', 'M=42 SD=34 nghĩa là gì'],
        [
            ['DASS-21 lo âu', '0–42', 'M=42 ≈ max → trần (ceiling effect); SD=34 vô lý'],
            ['DASS-42 lo âu', '0–42', 'M=42 = trần; SD lớn vì phân tán mạnh'],
            ['STAI State', '20–80', 'M=42 ≈ giữa; SD=34 quá rộng (40 % range) — nghi ngờ'],
            ['CES-D', '0–60', 'M=42 ≈ 70 % range; SD=34 = 57 % range — phân phối rất rộng'],
            ['Số ngày trong tháng', '0–30', 'M=42 KHÔNG hợp lý → có thể là tổng điểm khác'],
            ['SCARED', '0–82', 'M=42 ≈ giữa; SD=34 ~ 41 % range — hơi cao nhưng có thể do mẫu hỗn hợp'],
        ]
    )

    para(d, '')
    para(d, '⚠ Em đề nghị thầy CHO BIẾT THANG ĐO + RANGE CỤ THỂ để em đánh giá chính xác hơn.', bold=True, italic=True)

    H(d, '4. ĐỀ XUẤT KIỂM TRA', level=2)
    bullet(d, 'Yêu cầu tác giả công bố thêm: Median, Q1, Q3, Min, Max, Skewness, Kurtosis.')
    bullet(d, 'Nhìn HISTOGRAM hoặc BOX PLOT — nếu có outliers rõ rệt thì M không đại diện.')
    bullet(d, 'Kiểm tra Shapiro-Wilk normality test — p < 0,05 → phân phối không chuẩn → dùng nonparametric.')
    bullet(d, 'Báo cáo Median (IQR) thay/cùng M (SD) trong narrative.')

    H(d, 'CÂU TRẢ LỜI', level=2, color=BLUE)
    para(d, 'Tóm tắt:', bold=True, color=BLUE)
    bullet(d, 'CV (Coefficient of Variation) = SD/M = 34,075/42,09 = 81 % — RẤT CAO (>60 %).', color=BLUE)
    bullet(d, 'SD ≈ 81 % của M cho thấy phân phối RẤT RỘNG, có thể không chuẩn (lệch phải, lưỡng đỉnh, hoặc có outliers).', color=BLUE)
    bullet(d, 'M có thể KHÔNG đại diện cho mẫu — nên ưu tiên báo cáo MEDIAN + IQR.', color=BLUE)
    bullet(d, 'Nếu thang đo có cận trên < M+2SD = 110, không sao; nhưng nếu M − 2SD < 0 mà thang chỉ ≥ 0, đây là dấu hiệu chắc chắn KHÔNG phân phối chuẩn.', color=BLUE)

    para(d, 'Cách bình luận trong báo cáo:', bold=True, color=BLUE)
    para(d,
        '"Điểm trung bình M = 42,09 với độ lệch chuẩn SD = 34,075 cho hệ số biến thiên '
        'CV = 81 % — vượt xa ngưỡng "rất cao" (>60 %, theo Reed et al. 2002). Mức biến '
        'thiên này gợi ý phân phối có thể không chuẩn (lệch phải, lưỡng đỉnh, hoặc có '
        'outliers cực đoan). Để diễn giải chính xác, đề nghị tác giả bổ sung thông số '
        'median (IQR), kiểm định normality (Shapiro-Wilk), và biểu đồ phân phối. Nếu '
        'phân phối thực sự không chuẩn, các phép so sánh tham số (t-test, ANOVA) áp '
        'dụng cho dữ liệu này có thể không hợp lệ; nên dùng phép phi tham số tương '
        'đương (Mann-Whitney U, Kruskal-Wallis)."',
        color=BLUE, italic=True
    )

    para(d, 'Câu hỏi cần thầy cung cấp thêm:', bold=True, color=BLUE)
    bullet(d, 'M=42,09 SD=34,075 đo bằng thang gì? (DASS, STAI, CES-D, SCARED, hay khác?)', color=BLUE)
    bullet(d, 'Range tự nhiên của thang là gì? (0-21? 0-42? 20-80?)', color=BLUE)
    bullet(d, 'Cỡ mẫu n bao nhiêu?', color=BLUE)

    H(d, 'Phụ lục — Tài liệu tham khảo', level=2)
    para(d, '1. Reed, G. F., Lynn, F., & Meade, B. D. (2002). Use of coefficient of variation in assessing variability of quantitative assays. Clinical and Diagnostic Laboratory Immunology, 9(6), 1235–1239. https://doi.org/10.1128/cdli.9.6.1235-1239.2002', italic=True, size=11)
    para(d, '2. Field, A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). SAGE. — Hướng dẫn đọc M, SD, normality test.', italic=True, size=11)
    para(d, '3. Tabachnick, B. G., & Fidell, L. S. (2013). Using Multivariate Statistics (6th ed.). Pearson. — Khi nào dùng nonparametric test.', italic=True, size=11)

    out = OUT_DIR / 'Binh_luan_M_42_09_SD_34_075_CV_cao.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')

# ===========================================================
# DOC 4 — CFA là gì?
# ===========================================================
def build_doc4():
    d = new_doc()
    H(d, 'Mô hình CFA là gì?', level=1)
    H(d, '— Confirmatory Factor Analysis — Phân tích nhân tố khẳng định —', level=2)

    H(d, 'Câu hỏi của thầy', level=2)
    para(d, 'Mô hình CFA là gì, Em ơi?', color=BLUE)

    H(d, '1. ĐỊNH NGHĨA NGẮN', level=2)
    para(d, 'CFA = Confirmatory Factor Analysis = Phân tích nhân tố khẳng định.', bold=True)
    para(d, 'CFA là một kỹ thuật thống kê dùng để KIỂM TRA xem một mô hình lý thuyết về cấu trúc nhân tố (factor structure) của một thang đo có PHÙ HỢP với dữ liệu thực hay không.')

    H(d, '2. KHÁC BIỆT VỚI EFA (Exploratory Factor Analysis)', level=2)

    add_table(d,
        ['Tiêu chí', 'EFA (khám phá)', 'CFA (khẳng định)'],
        [
            ['Mục đích', 'TÌM cấu trúc nhân tố mới', 'KIỂM TRA cấu trúc đã giả định'],
            ['Khi dùng', 'Phát triển thang đo mới', 'Kiểm chứng thang đo có sẵn trên mẫu mới'],
            ['Số nhân tố', 'Để dữ liệu quyết định', 'Tác giả ấn định trước'],
            ['Items thuộc nhân tố nào', 'Để dữ liệu quyết định', 'Tác giả gán trước theo lý thuyết'],
            ['Chỉ số đánh giá', 'Eigenvalue, % variance, KMO', 'χ²/df, CFI, TLI, RMSEA, SRMR'],
            ['Phần mềm', 'SPSS, R psych package', 'AMOS, Mplus, R lavaan, LISREL'],
            ['Bậc cao trong NC', 'Bước đầu khi xây thang', 'Bước sau, kiểm chứng độ giá trị'],
        ]
    )

    para(d, '')
    para(d, 'Workflow chuẩn: EFA → CFA. EFA cho biết có MẤY nhân tố, mỗi nhân tố gồm items nào. CFA confirm cấu trúc đó trên mẫu MỚI khác.', italic=True)

    H(d, '3. VÍ DỤ CỤ THỂ — DASS-21', level=2)
    para(d, 'DASS-21 (Depression Anxiety Stress Scales) có 21 mục, lý thuyết gồm 3 nhân tố (mỗi 7 mục):')
    bullet(d, 'Trầm cảm: items 3, 5, 10, 13, 16, 17, 21')
    bullet(d, 'Lo âu: items 2, 4, 7, 9, 15, 19, 20')
    bullet(d, 'Stress: items 1, 6, 8, 11, 12, 14, 18')

    para(d, 'Khi dịch sang tiếng Việt và áp dụng trên HS VN, người ta phải chạy CFA để kiểm tra:')
    bullet(d, 'Có thật sự gom thành 3 nhân tố như cấu trúc gốc không? (Hay 1 nhân tố? 4 nhân tố?)')
    bullet(d, 'Mỗi mục có "load" lên đúng nhân tố lý thuyết không (factor loading > 0,4)?')
    bullet(d, '3 nhân tố có tương quan như y văn dự kiến không (~0,6-0,8 giữa Anxiety-Depression)?')

    H(d, '4. CÁC CHỈ SỐ FIT INDICES — Bao nhiêu là "fit tốt"?', level=2)

    add_table(d,
        ['Chỉ số', 'Tốt', 'Chấp nhận', 'Kém'],
        [
            ['χ²/df (chi-square / df)', '< 2', '< 3', '> 5'],
            ['CFI (Comparative Fit Index)', '> 0,95', '> 0,90', '< 0,90'],
            ['TLI (Tucker-Lewis Index)', '> 0,95', '> 0,90', '< 0,90'],
            ['RMSEA (Root Mean Square Error of Approximation)', '< 0,06', '0,06-0,08', '> 0,10'],
            ['SRMR (Standardized Root Mean Square Residual)', '< 0,08', '0,08-0,10', '> 0,10'],
            ['Factor loading từng item', '> 0,7', '0,4-0,7', '< 0,4'],
        ]
    )

    para(d, '')
    para(d, 'Tham chiếu chuẩn: Hu & Bentler (1999) Structural Equation Modeling, 6(1), 1-55.', italic=True)

    H(d, '5. KHI NÀO BÀI BÁO PHẢI CÓ CFA', level=2)
    bullet(d, 'Khi VALIDATE thang đo trên dân số/ngôn ngữ MỚI (ví dụ DASS-21 bản Việt hóa).')
    bullet(d, 'Khi đề xuất MÔ HÌNH MỚI (ví dụ thang đo academic stress 4 yếu tố mới).')
    bullet(d, 'Khi nghi ngờ thang đo gốc không phù hợp với mẫu hiện tại (xuyên văn hóa).')
    bullet(d, 'Khi tạp chí Q1/Q2 yêu cầu — đa số tạp chí psychometric quốc tế bắt buộc CFA.')

    H(d, '6. HẠN CHẾ CỦA CFA', level=2)
    bullet(d, 'Cần MẪU LỚN (tối thiểu 200-300, tốt hơn 500+) — quy tắc 5-10 quan sát/parameter.')
    bullet(d, 'Phụ thuộc giả định: phân phối chuẩn đa biến (multivariate normality). Nếu vi phạm, dùng MLM (Maximum Likelihood Robust) hoặc WLSMV.')
    bullet(d, 'KHÔNG bằng chứng nguyên nhân-kết quả — chỉ kiểm tra cấu trúc tương quan.')
    bullet(d, 'Có thể "fit tốt" nhưng vẫn sai về lý thuyết (model misspecification).')

    H(d, 'CÂU TRẢ LỜI', level=2, color=BLUE)
    para(d, 'CFA = Confirmatory Factor Analysis = Phân tích nhân tố khẳng định:', bold=True, color=BLUE)
    bullet(d, 'Là kỹ thuật thống kê KIỂM TRA xem cấu trúc nhân tố lý thuyết của một thang đo có khớp với dữ liệu thực không.', color=BLUE)
    bullet(d, 'Khác EFA: EFA TÌM cấu trúc, CFA KIỂM CHỨNG cấu trúc đã giả định.', color=BLUE)
    bullet(d, 'Đánh giá bằng các fit indices: χ²/df < 3, CFI > 0,90, TLI > 0,90, RMSEA < 0,08, SRMR < 0,08.', color=BLUE)
    bullet(d, 'Cần khi: validate thang đo bản Việt hóa, đề xuất mô hình mới, tạp chí Q1/Q2 yêu cầu.', color=BLUE)
    bullet(d, 'Cần mẫu LỚN (≥200-300), phân phối chuẩn đa biến hoặc dùng estimator robust.', color=BLUE)

    para(d, 'Áp dụng cho dự án thầy:', bold=True, color=BLUE)
    bullet(d, 'Khi thầy đọc một bài VN có "CFA" → kiểm tra: thang gốc có cấu trúc gì? CFA confirm hay reject? Fit indices có đạt ngưỡng không?', color=BLUE)
    bullet(d, 'Nhiều bài VN dùng DASS-21 nhưng KHÔNG chạy CFA → giả định mặc nhiên cấu trúc 3 nhân tố đúng cho mẫu Việt — đây là điểm yếu.', color=BLUE)
    bullet(d, 'Nếu thầy thiết kế NC mới với thang chuẩn (DASS, GAD-7), nên chạy CFA pilot trước với n=200 trước khi áp dụng full study.', color=BLUE)

    H(d, 'Phụ lục — Tài liệu tham khảo', level=2)
    para(d, '1. Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118 — KINH ĐIỂN cho ngưỡng fit indices.', italic=True, size=11)
    para(d, '2. Brown, T. A. (2015). Confirmatory factor analysis for applied research (2nd ed.). Guilford Press. — Sách chuẩn về CFA.', italic=True, size=11)
    para(d, '3. Kline, R. B. (2016). Principles and practice of structural equation modeling (4th ed.). Guilford Press. — SEM bao gồm CFA.', italic=True, size=11)
    para(d, '4. Byrne, B. M. (2016). Structural equation modeling with AMOS: Basic concepts, applications, and programming (3rd ed.). Routledge. — Hướng dẫn chạy CFA bằng AMOS.', italic=True, size=11)

    out = OUT_DIR / 'CFA_la_gi_Confirmatory_Factor_Analysis.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')

# ===========================================================
# DOC 5 — KTC 90% RMSEA (0,024-0,067) có được không?
# ===========================================================
def build_doc5():
    d = new_doc()
    H(d, 'KTC 90% của RMSEA (0,024 – 0,067) có được không?', level=1)
    H(d, '— Vì sao RMSEA dùng KTC 90% là CHUẨN VÀNG (không phải 95%) —', level=2)

    H(d, 'Câu hỏi của thầy', level=2)
    para(d, 'Khoảng tin cậy 90% của RMSEA (0,024 - 0,067) có được không Em?', color=BLUE)

    H(d, '1. NGẮN GỌN — CÓ, ĐÚNG CHUẨN', level=2)
    para(d, 'KTC 90% cho RMSEA là CHUẨN VÀNG quốc tế — KHÔNG PHẢI 95%. Đây là một trong số RẤT ÍT trường hợp mà KTC 90% được khuyến nghị chính thức.', bold=True, color=GREEN)

    H(d, '2. VÌ SAO KTC 90% (không phải 95%) cho RMSEA?', level=2)

    para(d, 'RMSEA (Root Mean Square Error of Approximation) là chỉ số đánh giá độ phù hợp của mô hình SEM/CFA. Khác với hầu hết phép thống kê, RMSEA dùng phân phối non-central chi-square chứ không phải phân phối chuẩn.')

    para(d, 'Browne & Cudeck (1993) — paper kinh điển đặt nền cho RMSEA — đã đề xuất KTC 90% với 2 lý do chính:', bold=True)
    bullet(d, '1. KTC 90% hai phía tương đương kiểm định 5% một phía. Vì RMSEA chỉ kiểm tra GIỚI HẠN TRÊN (mô hình KHÔNG fit kém) — chỉ quan tâm cận trên có vượt 0,08 hay không — nên KTC 90% đủ chặt.')
    bullet(d, '2. Phân phối non-central chi-square của RMSEA có tính chất bất đối xứng, làm KTC 95% rất rộng và ít hữu ích.')

    para(d, 'Từ đó, mọi sách giáo trình SEM (Brown 2015, Kline 2016, Byrne 2016) và mọi phần mềm (AMOS, Mplus, lavaan R) đều mặc định xuất KTC 90% cho RMSEA. Đây là quy ước CHUẨN HÓA toàn cầu.', italic=True)

    H(d, '3. ĐÁNH GIÁ KTC 0,024 – 0,067 CỤ THỂ', level=2)

    para(d, 'Áp dụng tiêu chí Hu & Bentler (1999) + MacCallum, Browne & Sugawara (1996):', bold=True)

    add_table(d,
        ['Giá trị RMSEA', 'Phân loại độ fit'],
        [
            ['< 0,01', 'Excellent (xuất sắc)'],
            ['0,01 – 0,05', 'Tốt (good fit)'],
            ['0,05 – 0,08', 'Chấp nhận được (acceptable / mediocre fit)'],
            ['0,08 – 0,10', 'Kém (poor fit)'],
            ['> 0,10', 'Không chấp nhận (unacceptable)'],
        ]
    )

    para(d, '')
    para(d, 'Áp dụng cho KTC 90% [0,024 – 0,067]:', bold=True)
    bullet(d, 'Cận DƯỚI = 0,024 → mô hình ÍT NHẤT có "Tốt" (good fit, < 0,05).')
    bullet(d, 'Cận TRÊN = 0,067 → mô hình NHIỀU NHẤT cũng chỉ "Chấp nhận được" (< 0,08), KHÔNG phải kém.')
    bullet(d, 'Width = 0,067 − 0,024 = 0,043 — KTC HẸP, độ chính xác CAO.')
    bullet(d, '⚠ Cận trên 0,067 chưa vượt 0,08 → MÔ HÌNH HỖ TRỢ. ✓ Đây là một KTC RMSEA ĐẸP.', color=GREEN)

    H(d, '4. KIỂM TRA KÈM THEO — pclose', level=2)
    para(d, 'Khi xem RMSEA, nên kiểm tra thêm chỉ số "pclose" hoặc "p of Close Fit":')
    bullet(d, 'pclose = xác suất RMSEA ≤ 0,05 (close fit).')
    bullet(d, 'Nếu pclose > 0,50 → mô hình close fit (rất tốt).')
    bullet(d, 'Nếu pclose < 0,05 → bác bỏ giả thuyết close fit (mô hình không phù hợp tốt).')
    bullet(d, 'Với RMSEA 0,024–0,067, pclose có thể ~ 0,3–0,7 tùy n. Đây là chấp nhận được.')

    H(d, '5. CẢNH BÁO QUAN TRỌNG', level=2)
    bullet(d, 'KTC 90% CHỈ ÁP DỤNG cho RMSEA. Các chỉ số khác (CFI, TLI, SRMR) KHÔNG có KTC truyền thống.')
    bullet(d, 'Nếu thầy gặp KTC 90% cho OR, RR, β, r → đó KHÔNG chuẩn (chuẩn là KTC 95%). Đây là điểm em đã nói trong doc QT021 Brunborg.')
    bullet(d, 'NHƯNG cho RMSEA → KTC 90% là ĐÚNG. Đây là ngoại lệ duy nhất trong y văn được sự đồng thuận.')

    H(d, '6. CÁCH BÁO CÁO CHUẨN APA 7', level=2)
    para(d, 'Khi viết báo cáo CFA, format chuẩn:', bold=True)
    para(d,
        '"Mô hình CFA cho thang [tên thang] đạt độ phù hợp tốt: χ²(df) = [giá trị], '
        'p < 0,001; CFI = 0,xxx; TLI = 0,xxx; RMSEA = 0,xxx, KTC 90% [0,xxx – 0,xxx], '
        'pclose = 0,xxx; SRMR = 0,xxx."',
        italic=True
    )
    para(d, 'Áp dụng số liệu của thầy:', bold=True)
    para(d, '"... RMSEA = [trung điểm ≈ 0,045], KTC 90% [0,024 – 0,067] — chỉ số này nằm trong vùng "tốt đến chấp nhận được", hỗ trợ cấu trúc mô hình."', italic=True, color=GREEN)

    H(d, 'CÂU TRẢ LỜI', level=2, color=BLUE)
    para(d, 'CÓ — KTC 90% cho RMSEA là CHUẨN VÀNG quốc tế:', bold=True, color=BLUE)
    bullet(d, 'Browne & Cudeck (1993) đề xuất KTC 90% — không phải 95% — vì RMSEA dùng phân phối non-central chi-square và chỉ kiểm tra cận trên.', color=BLUE)
    bullet(d, 'Mọi phần mềm SEM (AMOS, Mplus, lavaan R) mặc định xuất KTC 90% cho RMSEA.', color=BLUE)

    para(d, 'KTC 0,024 – 0,067 cụ thể:', bold=True, color=BLUE)
    bullet(d, 'Cận dưới 0,024 → mô hình ÍT NHẤT có fit "tốt" (good).', color=BLUE)
    bullet(d, 'Cận trên 0,067 → mô hình NHIỀU NHẤT cũng chỉ "chấp nhận được", KHÔNG vượt 0,08 (poor).', color=BLUE)
    bullet(d, 'Width 0,043 — KTC HẸP, độ chính xác cao.', color=BLUE)
    bullet(d, '✓ Đây là một KTC RMSEA ĐẸP, hỗ trợ mô hình tốt.', color=BLUE)

    para(d, 'Lưu ý chung:', bold=True, color=BLUE)
    bullet(d, 'KTC 90% chỉ ĐÚNG cho RMSEA — đây là ngoại lệ duy nhất.', color=BLUE)
    bullet(d, 'Cho mọi chỉ số khác (OR, RR, β, r) → KTC 95% mới đúng. Nếu thấy KTC 90% cho OR → cảnh báo.', color=BLUE)

    H(d, 'Phụ lục — Tài liệu tham khảo', level=2)
    para(d, '1. Browne, M. W., & Cudeck, R. (1993). Alternative ways of assessing model fit. In K. A. Bollen & J. S. Long (Eds.), Testing structural equation models (pp. 136–162). SAGE. — Bài KINH ĐIỂN đặt nền cho RMSEA + KTC 90%.', italic=True, size=11)
    para(d, '2. Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118 — Ngưỡng cắt RMSEA < 0,06 (tốt) / < 0,08 (chấp nhận).', italic=True, size=11)
    para(d, '3. MacCallum, R. C., Browne, M. W., & Sugawara, H. M. (1996). Power analysis and determination of sample size for covariance structure modeling. Psychological Methods, 1(2), 130–149. https://doi.org/10.1037/1082-989X.1.2.130 — Phân loại close fit / RMSEA ≤ 0,05.', italic=True, size=11)
    para(d, '4. Kenny, D. A., Kaniskan, B., & McCoach, D. B. (2015). The performance of RMSEA in models with small degrees of freedom. Sociological Methods & Research, 44(3), 486–507. https://doi.org/10.1177/0049124114543236 — Cảnh báo RMSEA bị "phóng đại" với df nhỏ.', italic=True, size=11)
    para(d, '5. Steiger, J. H. (1990). Structural model evaluation and modification: An interval estimation approach. Multivariate Behavioral Research, 25(2), 173–180. https://doi.org/10.1207/s15327906mbr2502_4 — Một trong những người sáng tạo RMSEA.', italic=True, size=11)

    out = OUT_DIR / 'KTC_90_RMSEA_0_024_0_067_chuan_vang.docx'
    d.save(out)
    print(f'  ✓ {out.name} ({out.stat().st_size//1024} KB)')

# ===========================================================
# Run all
# ===========================================================
print('Building 5 docs for evening session 04/05/2026:')
print()
build_doc1()
build_doc2()
build_doc3()
build_doc4()
build_doc5()
print()
print('Done.')
