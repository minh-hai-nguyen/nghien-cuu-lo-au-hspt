"""Build doc tra loi: Binh luan H.T.Hoc 2025 (VN014) ve 2 yeu to:
- Quan he cha me-con (β = 0,272)
- Thiet bi dien tu (β = 0,176)
Format: CAU TRA LOI to xanh truoc phu luc.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Binh_luan_VN014_HoangTrungHoc_2yeu_to_chamec_thietbi.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x30)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold

def para_black(text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(size); r.bold = bold; r.italic = italic

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)

def bullet_black(text, italic=False, size=12):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(size); r.italic = italic

def add_table(header, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

# =================================================================
H('Bình luận về 2 yếu tố trong nghiên cứu Hoàng Trung Học và cộng sự (2025)', level=1)
H('— Mối quan hệ cha mẹ-con & Sử dụng thiết bị điện tử —', level=2)

# CÂU HỎI
H('Câu hỏi của thầy', level=2)
para_blue(
    'Em cho thầy bình luận về H.T. Học và cs. (2025) với 2 yếu tố "mối quan hệ '
    'cha mẹ–con" và "sử dụng thiết bị điện tử" trong nghiên cứu về Mức độ căng '
    'thẳng, lo âu và trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 '
    'tại Việt Nam.'
)

# BỐI CẢNH
H('1. Bối cảnh nghiên cứu VN014', level=2)
para_black(
    'Nghiên cứu của Hoàng Trung Học và cộng sự (2025) — VN014 trong cơ sở dữ '
    'liệu dự án — là nghiên cứu cắt ngang tại hai thời điểm trên 8.389 thanh '
    'thiếu niên Việt Nam (Đợt 1 = 4.052 trong COVID 12/2021; Đợt 2 = 4.337 sau '
    'COVID 9/2023) tại 6 tỉnh. Công cụ chính là DASS-21 phiên bản tiếng Việt + '
    'bảng hỏi bổ sung về giấc ngủ, sử dụng thiết bị điện tử, hoạt động thể chất, '
    'và quan hệ gia đình. Hồi quy đa biến với R² = 0,190 (trong COVID).'
)

para_black('Bốn yếu tố ảnh hưởng được xác định và xếp hạng theo |β|:', bold=True)
add_table(
    ['Thứ', 'Yếu tố', 'Beta', 'Loại', 'Ý nghĩa'],
    [
        ['1', 'Mối quan hệ cha mẹ–con', '0,272', 'Nguy cơ', 'p < 0,001'],
        ['2', 'Thời gian dùng máy tính/điện thoại', '0,176', 'Nguy cơ', 'p < 0,001'],
        ['3', 'Giấc ngủ đủ', '–0,149', 'Bảo vệ', 'p < 0,001'],
        ['4', 'Hoạt động thể thao', '–0,087', 'Bảo vệ', 'p < 0,01'],
    ]
)
para_black('')
para_black('⚠ Lưu ý quan trọng về hướng β: Beta = 0,272 cho "mối quan hệ cha mẹ–con" mang dấu DƯƠNG, nghĩa là điểm cao của biến này tương quan với điểm DAS cao. Diễn giải này gợi ý biến đo MÂU THUẪN, CĂNG THẲNG, hoặc XUNG ĐỘT trong quan hệ — không phải chất lượng quan hệ tích cực. Bài gốc KHÔNG nêu rõ thang đo cụ thể (chỉ ghi "tính chất mối quan hệ với cha mẹ"), gây khó diễn giải chính xác. Đây là điểm yếu PHƯƠNG PHÁP đáng lưu ý khi trích dẫn.', bold=True, italic=True)

# YẾU TỐ 1
H('2. Yếu tố 1: Mối quan hệ cha mẹ–con (β = 0,272)', level=2)

H('2.1. Giá trị trong nghiên cứu', level=3)
para_black(
    'Beta chuẩn hóa 0,272 (p < 0,001) là HỆ SỐ MẠNH NHẤT trong bốn yếu tố ảnh '
    'hưởng được xác định. Theo phân loại Cohen (1988), |β| = 0,272 thuộc vùng '
    '"trung bình" (0,1 nhỏ; 0,3 trung bình; 0,5 lớn). Trong dân số mẫu lớn '
    '(n = 8.389), hệ số này có ý nghĩa thống kê chắc chắn và độ chính xác cao.'
)

H('2.2. Đối chiếu với corpus dự án', level=3)
add_table(
    ['Nghiên cứu', 'Beta/OR', 'Hướng', 'Diễn giải'],
    [
        ['VN014 Hoàng Trung Học (2025)', 'β = 0,272', 'Dương', 'Quan hệ cha mẹ-con (có thể mâu thuẫn) tăng → DAS tăng'],
        ['VN003 Phạm và cộng sự (2024) tại Huế', 'β = –0,40', 'Âm', 'Chăm sóc cảm xúc tăng → lo âu giảm'],
        ['QT009 Qiu và cộng sự (2022) Trung Quốc', 'OR = 0,30', 'Bảo vệ', 'Nuôi dạy tích cực giảm 70% trầm cảm'],
        ['QT009 Qiu và cộng sự (2022) Trung Quốc', 'OR = 2,01', 'Nguy cơ', 'Nuôi dạy tiêu cực tăng 2 lần lo âu'],
        ['V-NAMHS 2022 (UNICEF Việt Nam)', '5,1%', 'Khoảng trống', 'Phụ huynh xác định được con cần trợ giúp'],
    ]
)
para_black('')

para_black(
    'So sánh ba bài cho thấy bức tranh nhất quán — quan hệ gia đình là yếu tố '
    'TRUNG TÂM của lo âu vị thành niên Việt Nam và Á Châu. Tuy nhiên, ba bài '
    'đo các khía cạnh KHÁC NHAU của quan hệ này: VN014 đo "tính chất mối quan '
    'hệ" (có thể là mâu thuẫn); VN003 đo CHẤT LƯỢNG hỗ trợ cảm xúc; QT008 đo '
    'KIỂU NUÔI DẠY (positive vs negative parenting style). Phù hợp với khung '
    'của Pascoe, Hetrick và Parker (2020) trong International Journal of '
    'Adolescence and Youth — yếu tố gia đình tác động đa chiều lên lo âu '
    'thanh thiếu niên, không chỉ qua một biến đơn.'
)

H('2.3. Phản biện điểm yếu', level=3)
bullet_black('Bài KHÔNG nêu thang đo cụ thể cho "mối quan hệ cha mẹ–con". Chuẩn quốc tế: Parental Bonding Instrument (PBI; Parker và cộng sự, 1979), Family Assessment Device (FAD-12; Epstein và cộng sự, 1983), hoặc Parental Acceptance-Rejection Questionnaire (PARQ).')
bullet_black('Hướng β dương (0,272) gợi ý đo MÂU THUẪN hoặc THANG NGƯỢC — cần làm rõ trong báo cáo. Nếu đo "chất lượng quan hệ" theo chuẩn thông thường thì β phải ÂM (như VN003 Phạm 2024 với β = -0,40).')
bullet_black('Khoảng trống GIAO TIẾP gia đình–con không được đo. V-NAMHS 2022 ghi nhận chỉ 5,1% phụ huynh Việt Nam xác định được con cần trợ giúp tâm lý — nói cách khác, ngay cả khi cha mẹ có thiện chí, vẫn có khoảng trống nhận biết.')

H('2.4. Hàm ý can thiệp', level=3)
para_black(
    'Bất kể hướng β được diễn giải thế nào, kết quả khẳng định quan hệ cha mẹ–'
    'con là MỤC TIÊU CAN THIỆP ƯU TIÊN. Chương trình tập huấn cha mẹ '
    '(parenting training) đã được Cai và cộng sự (2025) trong phân tích tổng '
    'hợp các thử nghiệm ngẫu nhiên có đối chứng tại trường xác nhận hiệu quả '
    'tăng resilience và giảm triệu chứng tâm lý ở trẻ. Mô hình stepped-care '
    'của Matsumoto và cộng sự (2024) tại Nhật Bản đề xuất tích hợp tập huấn '
    'phụ huynh vào Tầng 2 (selective intervention) cho học sinh có yếu tố '
    'nguy cơ.'
)

# YẾU TỐ 2
H('3. Yếu tố 2: Sử dụng thiết bị điện tử (β = 0,176)', level=2)

H('3.1. Giá trị trong nghiên cứu', level=3)
para_black(
    'Beta chuẩn hóa 0,176 (p < 0,001) — yếu tố nguy cơ MẠNH THỨ HAI sau quan '
    'hệ cha mẹ–con. Theo Cohen (1988), |β| = 0,176 thuộc vùng "yếu–trung '
    'bình". Mặc dù không quá lớn, ý nghĩa thực hành đáng kể vì biến này có '
    'thể can thiệp được ở cả ba cấp độ — cá nhân, gia đình, và chính sách '
    'trường học.'
)

H('3.2. Đối chiếu với corpus dự án', level=3)
add_table(
    ['Nghiên cứu', 'Cường độ', 'Phát hiện chính'],
    [
        ['VN014 Hoàng Trung Học (2025)', 'β = 0,176', 'Thời gian dùng máy tính/điện thoại tăng → DAS tăng'],
        ['QT021 Brunborg và cộng sự (2025) Na Uy', 'β = 0,18 (nữ)', 'Mạng xã hội giải thích phần lớn xu hướng tăng distress'],
        ['QT033 Schmidt-Persson và cộng sự (2024) Đan Mạch', 'RCT n = 89', 'Hạn chế màn hình 14 ngày cải thiện vấn đề tâm lý nội hóa'],
        ['QT007 Chen và cộng sự (2023) Trung Quốc, BMC Psychiatry', 'OR = 5,00', 'Rối loạn chơi game tăng 5 lần lo âu'],
    ]
)
para_black('')

para_black(
    'Phù hợp với phát hiện toàn cầu — sử dụng thiết bị điện tử là yếu tố '
    'nguy cơ ĐƯỢC XÁC LẬP cho lo âu vị thành niên. Đáng chú ý, Brunborg và '
    'cộng sự (2025) trên 979.043 học sinh Na Uy ghi nhận β chuẩn hóa cho '
    'thời gian sử dụng mạng xã hội ở học sinh nữ là 0,18 — gần ngang với '
    'β của VN014. Nói cách khác, hai nghiên cứu trên hai nền văn hóa khác '
    'biệt ghi nhận cường độ tác động TƯƠNG TỰ — củng cố tính phổ quát của '
    'yếu tố này.'
)

H('3.3. Phản biện điểm yếu', level=3)
bullet_black('Bài CHỈ đo TẦN SUẤT (thời gian) sử dụng, KHÔNG đo NỘI DUNG hoặc CHẤT LƯỢNG. Brunborg và cộng sự (2025) cũng chỉ ra hạn chế tương tự — đo thời gian tổng (≥ 3 giờ/ngày) không phân biệt được hoạt động học tập, giải trí, hay tương tác xã hội.')
bullet_black('Không phân biệt LOẠI thiết bị (máy tính học tập, smartphone giải trí, máy tính bảng) hay LOẠI ứng dụng (mạng xã hội, video, game, học tập). Theo Schmidt-Persson và cộng sự (2024), hạn chế CỤ THỂ "màn hình giải trí" mới có hiệu quả cải thiện tâm lý — không phải hạn chế tổng.')
bullet_black('Không phân tích tương tác với GIỚI TÍNH. Brunborg 2025 cho thấy mạng xã hội tác động MẠNH HƠN đến học sinh nữ; VN014 không kiểm chứng được giả thuyết này.')
bullet_black('Cỡ mẫu n = 8.389 đủ lớn để phát hiện hiệu ứng tinh tế, nhưng do không có thiết kế dọc thực sự (chỉ hai cắt ngang trên hai mẫu khác nhau), không thể kết luận NHÂN QUẢ — chỉ là tương quan.')

H('3.4. Hàm ý can thiệp', level=3)
para_black(
    'Phù hợp với hai chiến lược can thiệp đã có bằng chứng. Thứ nhất, hạn '
    'chế thời gian màn hình GIẢI TRÍ trong 14 ngày như Schmidt-Persson và '
    'cộng sự (2024) đã chứng minh hiệu quả trong thử nghiệm SCREENS-Kids. '
    'Thứ hai, giáo dục digital wellbeing cho học sinh — phát triển kỹ năng '
    'tự điều chỉnh thời gian, nhận biết tác động tâm lý của mạng xã hội, và '
    'tập trung vào nội dung tích cực. Cả hai chiến lược cần được tích hợp '
    'với sự ĐỒNG HÀNH của phụ huynh — kết nối với Yếu tố 1.'
)

# SO SÁNH HAI YẾU TỐ
H('4. So sánh hai yếu tố và hàm ý cho can thiệp', level=2)

para_black(
    'Cả hai yếu tố là YẾU TỐ NGUY CƠ có thể can thiệp được — khác với các '
    'yếu tố nhân khẩu cố định như giới tính hoặc khối lớp. Tuy nhiên, cường '
    'độ tác động khác biệt rõ ràng — quan hệ cha mẹ–con (β = 0,272) mạnh '
    'hơn thiết bị điện tử (β = 0,176) khoảng 1,55 lần.'
)

add_table(
    ['Khía cạnh', 'Cha mẹ–con (β = 0,272)', 'Thiết bị điện tử (β = 0,176)'],
    [
        ['Cường độ', 'Mạnh nhất trong 4 yếu tố', 'Mạnh thứ hai'],
        ['Cấp độ can thiệp', 'Gia đình + cá nhân', 'Cá nhân + chính sách trường'],
        ['Thời gian thấy hiệu quả', 'Trung-dài hạn (3-12 tháng)', 'Ngắn hạn (2 tuần — Schmidt-Persson 2024)'],
        ['Chi phí can thiệp', 'Cao (cần tập huấn cha mẹ)', 'Thấp-vừa (giáo dục digital wellbeing)'],
        ['Khả năng nhân rộng', 'Khó (cần thay đổi văn hóa gia đình)', 'Dễ (qua trường học)'],
        ['Khuyến nghị ưu tiên', 'Tầng 2-3 (selective + indicated)', 'Tầng 1-2 (universal + selective)'],
    ]
)
para_black('')

para_black(
    'Phát hiện đáng lưu ý — hai yếu tố KHÔNG độc lập với nhau. Khi quan hệ '
    'cha mẹ–con bị căng thẳng, học sinh thường rút vào thiết bị điện tử như '
    'một cách trốn thoát hoặc tìm hỗ trợ thay thế (peer online). Phù hợp '
    'với mô hình của Pascoe và cộng sự (2020) về tương tác giữa các yếu tố '
    'gia đình và hành vi sức khỏe. Đề xuất nghiên cứu tiếp theo — kiểm '
    'chứng hiệu ứng trung gian của thiết bị điện tử trong mối quan hệ giữa '
    'mâu thuẫn gia đình và rối loạn lo âu.'
)

# CÂU TRẢ LỜI
H('5. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Tóm gọn về hai yếu tố:', bold=True)

bullet_blue(
    'Yếu tố 1 — MỐI QUAN HỆ CHA MẸ–CON (β = 0,272; p < 0,001) là yếu tố '
    'ảnh hưởng MẠNH NHẤT trong nghiên cứu, cường độ "trung bình" theo Cohen '
    '(1988). Hướng β dương gợi ý biến đo MÂU THUẪN/CĂNG THẲNG, không phải '
    'chất lượng tích cực. Phù hợp với phát hiện của Phạm và cộng sự (2024) '
    'tại Huế (β = -0,40 cho chăm sóc cảm xúc) và Qiu và cộng sự (2022) tại '
    'Trung Quốc (OR = 2,01 cho nuôi dạy tiêu cực; OR = 0,30 cho nuôi dạy '
    'tích cực).'
)

bullet_blue(
    'Yếu tố 2 — SỬ DỤNG THIẾT BỊ ĐIỆN TỬ (β = 0,176; p < 0,001) là yếu tố '
    'nguy cơ MẠNH THỨ HAI, cường độ "yếu–trung bình". Cường độ này NGANG '
    'HÀNG với phát hiện của Brunborg và cộng sự (2025) tại Na Uy trên '
    '979.043 học sinh (β = 0,18 cho nữ) — gợi ý tính phổ quát xuyên văn '
    'hóa.'
)

para_blue('Ba điểm yếu phương pháp đáng phản biện:', bold=True)
bullet_blue(
    'Bài KHÔNG nêu thang đo cụ thể cho "mối quan hệ cha mẹ–con" — không '
    'rõ là Parental Bonding Instrument, FAD-12, hay tự thiết kế. Hướng β '
    'dương cần được làm rõ trong báo cáo.'
)
bullet_blue(
    'Bài CHỈ đo TẦN SUẤT thời gian thiết bị, KHÔNG đo NỘI DUNG hay LOẠI '
    'ứng dụng. Schmidt-Persson và cộng sự (2024) cho thấy chỉ "màn hình '
    'giải trí" có hiệu ứng tâm lý, không phải tổng thời gian.'
)
bullet_blue(
    'Thiết kế cắt ngang trên HAI MẪU KHÁC NHAU không cho phép kết luận '
    'NHÂN QUẢ — chỉ là tương quan. Cần nghiên cứu thuần tập (longitudinal) '
    'theo cùng một nhóm học sinh để xác lập quan hệ nhân quả.'
)

para_blue('Hàm ý can thiệp:', bold=True)
bullet_blue(
    'Quan hệ cha mẹ–con là MỤC TIÊU ƯU TIÊN — cần chương trình tập huấn '
    'phụ huynh (parenting training) ở Tầng 2-3 trong mô hình stepped-care '
    'của Matsumoto và cộng sự (2024). Phù hợp với phát hiện của Cai và '
    'cộng sự (2025) — can thiệp resilience kết hợp tập huấn cha mẹ tăng '
    'hiệu quả tổng thể.'
)
bullet_blue(
    'Thiết bị điện tử nên can thiệp ở Tầng 1 (universal) — giáo dục '
    'digital wellbeing toàn trường, kết hợp Tầng 2 hạn chế thời gian màn '
    'hình giải trí cho nhóm có yếu tố nguy cơ. Cần ĐỒNG HÀNH của phụ '
    'huynh để hiệu quả bền vững.'
)

para_blue('Khuyến nghị cách trích vào báo cáo CTH:', bold=True)
para_blue(
    '"Hoàng Trung Học và cộng sự (2025) trên 8.389 thanh thiếu niên 6 tỉnh '
    'Việt Nam xác định mối quan hệ cha mẹ–con (β = 0,272) và sử dụng thiết '
    'bị điện tử (β = 0,176) là hai yếu tố nguy cơ mạnh nhất ảnh hưởng đến '
    'mức độ trầm cảm, lo âu, căng thẳng (DAS). Cường độ tương đương với các '
    'phát hiện quốc tế (Brunborg và cộng sự, 2025; Phạm và cộng sự, 2024) '
    'gợi ý hai yếu tố này là MỤC TIÊU CAN THIỆP ƯU TIÊN trong chương trình '
    'sức khỏe tâm thần học đường tại Việt Nam, với hướng tiếp cận stepped-'
    'care theo mô hình của Matsumoto và cộng sự (2024)."'
)

# PHỤ LỤC
H('6. Phụ lục — Tài liệu tham khảo', level=2)

para_black('Tiếng Việt', bold=True)
para_black(
    '1. Hoàng Trung Học và cộng sự. (2025). Mức độ căng thẳng, lo âu và '
    'trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 tại Việt '
    'Nam: Nghiên cứu cắt ngang. Asian Journal of Public Research. [VN014 '
    'trong cơ sở dữ liệu dự án.]', italic=True, size=11
)
para_black(
    '2. Phạm, V. T. và cộng sự. (2024). Mối liên hệ giữa hỗ trợ xã hội và '
    'sức khỏe tâm thần ở thanh thiếu niên tại Huế, Việt Nam. [VN003 trong '
    'cơ sở dữ liệu dự án.]', italic=True, size=11
)
para_black(
    '3. UNICEF Việt Nam, Bộ Lao động – Thương binh và Xã hội, và Tổng '
    'cục Thống kê. (2022). Khảo sát Sức khỏe Tâm thần Vị thành niên Việt '
    'Nam (V-NAMHS 2022). Hà Nội. [VN002 trong cơ sở dữ liệu dự án.]',
    italic=True, size=11
)

para_black('Tiếng Anh', bold=True)
para_black(
    '4. Brunborg, G. S., Nilsen, S. A., Skogen, J. C., & Bang, L. (2025). '
    'Possible explanations for the upward trend in mental distress among '
    'adolescents in Norway from 2011 to 2024. Social Science & Medicine, '
    '384, 118528. https://doi.org/10.1016/j.socscimed.2025.118528 '
    '[QT021 trong cơ sở dữ liệu dự án.]', italic=True, size=11
)
para_black(
    '5. Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based '
    'interventions for resilience in children and adolescents: A '
    'systematic review and meta-analysis of randomized controlled trials. '
    'Frontiers in Psychiatry, 16, 1594658. [QT044 trong cơ sở dữ liệu '
    'dự án.]', italic=True, size=11
)
para_black(
    '6. Chen, Z., et al. (2023). Bullying, sleep, and gaming disorder among '
    'Chinese adolescents in Western China. BMC Psychiatry, 23, 580. '
    '[QT007 trong cơ sở dữ liệu dự án.]', italic=True, size=11
)
para_black(
    '7. Cohen, J. (1988). Statistical power analysis for the behavioral '
    'sciences (2nd ed.). Lawrence Erlbaum Associates.', italic=True, size=11
)
para_black(
    '8. Epstein, N. B., Baldwin, L. M., & Bishop, D. S. (1983). The '
    'McMaster Family Assessment Device. Journal of Marital and Family '
    'Therapy, 9(2), 171–180. https://doi.org/10.1111/j.1752-0606.1983.tb01497.x',
    italic=True, size=11
)
para_black(
    '9. Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the '
    'Depression Anxiety Stress Scales (2nd ed.). Sydney: Psychology '
    'Foundation of Australia.', italic=True, size=11
)
para_black(
    '10. Matsumoto, K., et al. (2024). Internet-based cognitive '
    'behavioral therapy for Japanese adolescents with anxiety and '
    'depression. JMIR Mental Health. [QT045 trong cơ sở dữ liệu dự án.]',
    italic=True, size=11
)
para_black(
    '11. Parker, G., Tupling, H., & Brown, L. B. (1979). A Parental '
    'Bonding Instrument. British Journal of Medical Psychology, 52(1), '
    '1–10. https://doi.org/10.1111/j.2044-8341.1979.tb02487.x', italic=True, size=11
)
para_black(
    '12. Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The '
    'impact of stress on students in secondary school and higher '
    'education. International Journal of Adolescence and Youth, 25(1), '
    '104–112. https://doi.org/10.1080/02673843.2019.1596823 [QT067 trong '
    'cơ sở dữ liệu dự án.]', italic=True, size=11
)
para_black(
    '13. Qiu, Z., Guo, Y., Wang, J., & Zhang, H. (2022). Associations of '
    'parenting style and resilience with depression and anxiety among '
    'Chinese high school students. Frontiers in Public Health, 10, '
    '989125. https://doi.org/10.3389/fpubh.2022.989125 [QT009 trong cơ '
    'sở dữ liệu dự án.]', italic=True, size=11
)
para_black(
    '14. Schmidt-Persson, J., et al. (2024). Screen media use and mental '
    'health of children and adolescents: A secondary analysis of the '
    'SCREENS-Kids randomized clinical trial. JAMA Network Open, 7(1), '
    'e2354033. [QT033 trong cơ sở dữ liệu dự án.]', italic=True, size=11
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
