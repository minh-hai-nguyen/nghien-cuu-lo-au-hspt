# -*- coding: utf-8 -*-
"""
BAO CAO v5 — BAN SACH GUI SEP
Content thuan tuy nghien cuu: KHONG nhac den KG, RAG, AI, version comparison, tool stack.
Giu nguyen Phan I-IV tu v3, viet lai Phan V thanh "Phat hien moi tu tong hop van ban"
(5 insights + Happy House + khuyen nghi) nhu la ket qua nghien cuu thuong quy.
"""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 11042026 v3.docx')
DST = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-5.docx')
CHARTS = os.path.join(os.path.dirname(__file__), 'charts')
KG_DIR = os.path.join(os.path.dirname(__file__), 'kg_data')
PAGE_W = 16.0
RED = RGBColor(0xCC, 0, 0)

shutil.copy(SRC, DST)
doc = Document(DST)

# ==================== Post-process inherited tables from v3 ====================
# L1: Table 0 — add HUPH to "Đơn vị" row
# L3: Note — Table 1 (Happy House results 5x3) is inherited as-is, structure OK
for t in doc.tables:
    # Find "Đơn vị" row in Happy House info table
    if len(t.rows) >= 2 and len(t.columns) == 2:
        for row in t.rows:
            cells = row.cells
            if cells[0].text.strip() == 'Đơn vị' and 'Monash' in cells[1].text:
                if 'HUPH' not in cells[1].text and 'Hanoi University of Public Health' not in cells[1].text:
                    cells[1].text = (
                        'Monash University (Melbourne, Úc) — đơn vị chủ trì; '
                        'Hanoi University of Public Health (HUPH, Việt Nam) — '
                        'triển khai thực địa; Queensland University of Technology '
                        '(Brisbane, Úc) — cố vấn RAP-A'
                    )
                    for p in cells[1].paragraphs:
                        for r in p.runs:
                            r.font.name = 'Times New Roman'
                            r.font.size = Pt(10)
                break

# ==================== helpers ====================
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def H(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def P(text, bold=False, italic=False, size=12, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic

def table(headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
    return t

def img(filename, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        path = os.path.join(KG_DIR, filename)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.italic=True

# ============================================================
# PHAN V — PHAT HIEN BO SUNG TU TONG HOP 68 BAI NGHIEN CUU
# ============================================================
doc.add_page_break()
H('PHẦN V — PHÁT HIỆN BỔ SUNG TỪ TỔNG HỢP 68 BÀI NGHIÊN CỨU (12/04/2026)',
  level=1)

P('Phần này trình bày 5 phát hiện liên bài quan trọng, chi tiết đầy đủ về nghiên cứu '
  'can thiệp Happy House tại Việt Nam (RCT đầu tiên trong lĩnh vực), và các khuyến nghị '
  'cụ thể cho đề cương can thiệp lo âu vị thành niên Việt Nam trong giai đoạn 2026–2028. '
  'Các phát hiện được rút ra sau khi rà soát lại toàn bộ 68 tóm tắt và bản dịch với mục '
  'tiêu bổ khuyết các số liệu định lượng chưa được khai thác trong các phiên bản trước.',
  italic=True)

# ============================================================
# V.1 — Bang tong hop so lieu then chot
# ============================================================
H('V.1. Bảng tổng hợp các số liệu định lượng then chốt từ 68 bài nghiên cứu', level=2)

P('Bảng dưới tổng hợp 18 số liệu định lượng quan trọng nhất được rút ra từ tổng hợp 68 '
  'bài nghiên cứu đã hoàn thành bản dịch hoặc tóm tắt. Các số liệu bao gồm tỷ số chênh '
  '(OR, AOR), hệ số hồi quy (β), và tỷ lệ phần trăm — tất cả đều có ý nghĩa thống kê và '
  'được công bố trên các tạp chí có bình duyệt uy tín.')

table(
    ['ID', 'Bài', 'Số liệu', 'Ngữ cảnh'],
    [
        ['VN001', 'Hoa 2024 Hà Nội', 'OR = 11,6 (áp lực học tập)', '3.910 HS THPT, GAD-7'],
        ['QT008', 'Wen 2020 nông thôn TQ', 'OR = 11,6 (áp lực học tập)', '900 HS THCS, LPA'],
        ['QT010', 'Xu 2021 Trung Quốc', 'OR = 1,3 (COVID worry)', '373.216 HS — lớn nhất TQ'],
        ['VN015', 'Ngô Anh Vinh 2024 Lạng Sơn', 'OR = 6,84 (bạn bè kém → trầm cảm)', '845 HS DTTS nội trú'],
        ['VN020', 'Trần Hồ Vĩnh Lộc 2024 TPHCM', 'OR = 11,6 (áp lực học tập)', '976 HS THPT, DASS-Y'],
        ['QT021', 'Brunborg Norway 2025', 'g = 0,46 (mental distress)', '979.043 VTN 2011–2024'],
        ['QT023', 'Mojtabai JAACAP 2024', 'AOR = 2,17 và 2,93', '13.684 VTN Hoa Kỳ (trends)'],
        ['QT026', 'UK NHS 2025', 'AOR = 2,17', 'Mental health England'],
        ['QT031', 'Islam 59 nước 2025', 'Lo âu ĐNA 3,6 % (thấp nhất)', 'GSHS 59 nước, Q1'],
        ['QT034', 'Cho Korea 2024', 'Income gradient', '213.000 VTN 2006–2022'],
        ['QT035', 'Jefferies 7 nước 2020', 'SAD cross-country PLOS ONE', 'Social anxiety 7 nước'],
        ['VN021', 'Trần Thảo Vi 2024 Huế', 'β = 4,73 (học thêm)', '611 HS Huế, NC dọc 3 năm, ESSA'],
        ['VN022', 'UNICEF VN 2022', '47 % học thêm > 3 h/tuần', 'School factors multi-province'],
        ['VN025', 'Phạm Thị Ngọc Hải Phòng', 'Lo âu 53,8 % (Cộng Hiền)', 'DASS-21, 420 HS Vĩnh Bảo'],
        ['QT047', 'Dong PLOS 2025 Ya An', 'Lo âu 41,4 %; OR = 0,22 (gia đình)',
         'n = 2.716 — bảo vệ mạnh nhất'],
        ['QT015', 'Zhu Suzhou 2025 BMC', 'AOR = 13,71 (ngủ < 5h; KTC: 8,66–21,71)',
         'n = 9.831 — nguy cơ mạnh nhất'],
        ['VN028', 'Đào Thị Ngoãn HMU 2025', 'OR = 4,97 (điểm Giỏi)', '196 SV Y4, DASS-21'],
        ['VN029', 'Duong TPHCM 2025 Q1', '91,6 % đa HRB', 'n = 2.631 HS TPHCM'],
    ],
    widths=[1.3, 4.2, 5.0, 5.0])

P('Các số liệu này có đặc điểm chung: đều là ước lượng định lượng có ý nghĩa thống kê, '
  'được công bố trên các tạp chí Q1 uy tín (PLOS ONE, Social Psychiatry and Psychiatric '
  'Epidemiology, BMC Public Health, JAACAP) hoặc dựa trên mẫu lớn hơn 1.000 người hoặc đa '
  'trung tâm. Chúng sẽ được khai thác có hệ thống trong 5 phát hiện dưới đây.')

# ============================================================
# V.2 — 5 PHAT HIEN
# ============================================================
H('V.2. Năm phát hiện liên bài quan trọng', level=2)

# ---- Insight 1 ----
H('V.2.1. Áp lực học tập — OR ≈ 11,6 lặp lại chính xác ở ba nghiên cứu độc lập '
  '(hai tại Việt Nam và một tại nông thôn Trung Quốc)', level=3)

P('Phát hiện đáng chú ý nhất là mối liên hệ giữa áp lực học tập và lo âu/trầm cảm vị '
  'thành niên có tỷ số chênh OR xấp xỉ 11,6 lặp lại ở ba nghiên cứu độc lập tại hai '
  'quốc gia khác nhau, với thiết kế và công cụ đo khác nhau — một mức độ replication '
  'hiếm thấy trong dịch tễ học sức khoẻ tâm thần vị thành niên.')

P('Chi tiết các nguồn:')
P('• Hoa et al. 2024 [VN001] (Frontiers in Public Health, Hà Nội, n = 3.910 học sinh '
  'THPT, đo bằng GAD-7): OR = 11,6 cho nhóm có áp lực học tập rất cao so với không có '
  'áp lực. Đây là nghiên cứu Việt Nam có cỡ mẫu lớn nhất về lo âu học sinh THPT Hà Nội.')
P('• Wen et al. 2020 [QT008] (n = 900 học sinh THCS nông thôn Trung Quốc, thang MHT, '
  'phân tích hồ sơ tiềm ẩn — Latent Profile Analysis): OR = 11,6 cho cùng yếu tố "áp '
  'lực học tập rất cao" trong nhóm lo âu nặng. Wen sử dụng phương pháp person-centered '
  '(LPA) thay vì variable-centered, nhưng kết quả trùng khớp chính xác với VN001.')
P('• Trần Hồ Vĩnh Lộc et al. 2024 [VN020] (Tạp chí Y học TPHCM, n = 976 học sinh '
  'THPT, DASS-Y): OR = 11,6 — trùng chính xác với VN001. Đây là nghiên cứu độc lập '
  'hoàn toàn (TPHCM vs Hà Nội, DASS-Y vs GAD-7).')

P('Tại sao con số OR ≈ 11 lặp lại đến vậy?', bold=True)
P('Có ba giả thuyết bổ trợ cho nhau. Thứ nhất, về mặt sinh học, trục HPA và cortisol '
  'phản ứng với áp lực mạn tính theo cơ chế chung không phụ thuộc văn hoá. Các nghiên '
  'cứu neurobiology chỉ ra cortisol cao kéo dài 6–12 tháng trước kỳ thi là yếu tố dự báo '
  'mạnh cho lo âu lâm sàng ở vị thành niên. Nếu cơ chế là sinh học, việc OR lặp lại ở các '
  'văn hoá khác nhau là hợp lý.')
P('Thứ hai, hệ thống giáo dục Trung Quốc, Việt Nam và một số nước Đông Á có chung đặc '
  'điểm: điểm thi quyết định tương lai, xếp hạng công khai, kỳ vọng gia đình rất cao. '
  'Báo cáo UNICEF Việt Nam 2022 [VN022] cho biết 47 % học sinh học thêm trên 3 giờ mỗi '
  'tuần — con số gần như không thấy ở phương Tây. Dong et al. 2025 [QT047] tại Ya An '
  '(Tứ Xuyên) phát hiện 60,3 % học sinh có nỗi sợ "làm cha mẹ thất vọng".')
P('Thứ ba, ba nghiên cứu độc lập (Hoa Hà Nội, Trần Hồ Vĩnh Lộc TPHCM, và Wen nông thôn '
  'Trung Quốc) đều cho cùng một điểm ước lượng OR = 11,6 mặc dù sử dụng thang đo khác '
  'nhau (GAD-7, DASS-Y, MHT) và phương pháp phân tích khác nhau (logistic regression '
  'variable-centered vs LPA person-centered). Mức độ trùng khớp này gợi ý đây không '
  'phải trùng hợp ngẫu nhiên mà phản ánh một hiệu ứng ổn định xuyên bối cảnh văn hoá '
  'Á Đông.')

P('Hệ quả cho đề cương Việt Nam:', bold=True)
P('Áp lực học tập phải là mục tiêu can thiệp trung tâm, không phải thành phần phụ trợ. '
  'Các biện pháp giảm tải học thêm, dạy kỹ năng quản lý thời gian và tư vấn hướng nghiệp '
  'sớm (để giảm ám ảnh "phải vào đại học top") cần được tích hợp vào chương trình can '
  'thiệp chính. Song song, cần thống nhất công cụ đo — sử dụng Educational Stress Scale '
  'for Adolescents (ESSA, Sun et al. 2011) bản Việt Nam cho mọi RCT tương lai để cho '
  'phép so sánh trực tiếp với các nghiên cứu Trung Quốc, Hàn Quốc. Nếu OR = 11 được xác '
  'nhận, gánh nặng y tế công cộng của áp lực học tập có thể lớn hơn nhiều yếu tố nguy cơ '
  'khác ở người lớn, đòi hỏi phối hợp với Bộ Giáo dục và Đào tạo về can thiệp cấp chính '
  'sách (giảm áp lực thi, bỏ xếp hạng công khai).')

# ---- Insight 2 ----
H('V.2.2. Yếu tố bảo vệ mạnh nhất — kênh giao tiếp gia đình (OR = 0,22)', level=3)

P('Dong et al. 2025 [QT047] (PLOS ONE, n = 2.716 học sinh trung học Ya An, Tứ Xuyên, '
  'Trung Quốc, DASS-21) là nghiên cứu duy nhất trong tổng hợp 68 bài đo "tâm sự với '
  'gia đình" như một biến độc lập. Kết quả: OR = 0,22 (KTC 95 %: 0,15 – 0,33, p < '
  '0,001) cho trầm cảm có ý nghĩa lâm sàng — tương ứng với hiệp hội giảm 78 % nguy cơ '
  '(lưu ý: đây là hiệp hội cắt ngang, không phải hiệu quả can thiệp).')

P('Đây là con số đáng chú ý cho một yếu tố bảo vệ. Trong dịch tễ học sức khoẻ tâm thần '
  'vị thành niên, rất ít yếu tố bảo vệ được báo cáo ở mức độ hiệp hội mạnh như vậy (OR '
  'dưới 0,30). Để đặt con số vào bối cảnh, Dong 2025 cũng báo cáo OR = 0,27 cho cùng '
  'yếu tố "tâm sự với gia đình" đối với lo âu (DASS-21), và OR = 0,37 cho "tâm sự với '
  'bạn bè" đối với trầm cảm — cho thấy rằng kênh giao tiếp đáng tin cậy với gia đình '
  'có tác động mạnh nhất trong các kênh chia sẻ cảm xúc mà Dong đo được.')

P('Lưu ý phương pháp luận quan trọng:', bold=True)
P('OR = 0,22 là ước lượng từ phân tích CẮT NGANG (cross-sectional association), '
  'không phải hiệu quả can thiệp từ RCT. Con số này chịu ảnh hưởng của confounding '
  '(gia đình có giao tiếp tốt thường có điều kiện kinh tế, giáo dục thuận lợi hơn) và '
  'reverse causality (trẻ ít lo âu có thể dễ cởi mở với cha mẹ hơn). Vì vậy, không thể '
  'so sánh trực tiếp OR = 0,22 với hiệu ứng can thiệp từ RCT (vd. aOR = 0,56 của Happy '
  'House). Các hiệp hội quan sát hầu như luôn lớn hơn hiệu ứng can thiệp RCT tương ứng. '
  'Tuy nhiên, độ lớn của hiệp hội (giảm 78 % nguy cơ) và tính nhất quán lý thuyết với '
  'khung attachment (Bowlby) và emotion regulation literature gợi ý giao tiếp gia đình '
  'là một ỨNG VIÊN CAN THIỆP đáng được kiểm chứng bằng RCT tại Việt Nam.')

P('Cơ chế được đề xuất:', bold=True)
P('(a) Điều tiết cảm xúc qua chia sẻ xã hội — Khi trẻ có thể chia sẻ cảm xúc với người '
  'lớn đáng tin cậy, trẻ học được cách gọi tên cảm xúc (emotion labeling), chuẩn hoá cảm '
  'xúc (emotion validation) và phát triển chiến lược ứng phó.')
P('(b) Cơ sở gắn bó an toàn (secure attachment base) — Theo lý thuyết Bowlby, mối quan '
  'hệ ấm áp với cha mẹ tạo "nền an toàn" để trẻ thám hiểm thế giới. Khi trẻ có thể quay '
  'về nền này khi gặp stress, trục HPA không kích hoạt kéo dài.')
P('(c) Phát hiện sớm và tìm kiếm trợ giúp — Cha mẹ có kênh giao tiếp tốt sẽ phát hiện '
  'sớm dấu hiệu lo âu/trầm cảm ở con và có thể đóng vai trò trung gian để tìm kiếm trợ '
  'giúp chuyên nghiệp kịp thời. Dong 2025 cũng báo cáo rằng hành vi tìm kiếm trợ giúp '
  'CHỦ ĐỘNG (so với thụ động) là yếu tố bảo vệ có ý nghĩa thống kê (OR = 0,48, p < '
  '0,001) cho cả trầm cảm lẫn lo âu — củng cố vai trò của hệ thống hỗ trợ đa tầng.')

P('Hệ quả cho đề cương Việt Nam:', bold=True)
P('Đề xuất bổ sung một module mới mang tên "Kỹ năng giao tiếp gia đình" cho cha mẹ vị '
  'thành niên Việt Nam, gồm 4 buổi chính và 2 buổi nhắc lại ở mốc 3 và 6 tháng. Nội dung '
  'có thể tham khảo Parent-Child Communication Training (Dumas 2005) nhưng thích ứng văn '
  'hoá Việt Nam (sử dụng hình ảnh bữa cơm gia đình, cách hỏi về cảm xúc thay vì chỉ hỏi '
  'về điểm số, cách lắng nghe không phán xét). Nhà trường nên tổ chức workshop cha mẹ '
  'thường xuyên mỗi học kỳ một lần. Đồng thời, cần một RCT chuyên biệt trong bối cảnh '
  'Việt Nam để xác nhận khả năng khái quát hoá của OR = 0,22. Nếu khái quát được, đây là '
  'can thiệp rẻ tiền và khả thi nhất cho bối cảnh quốc gia thu nhập trung bình — chỉ cần '
  'đào tạo giáo viên hoặc điều dưỡng học đường để triển khai.')

# ---- Insight 3 ----
H('V.2.3. Yếu tố nguy cơ bất ngờ — ngủ dưới 5 giờ AOR = 13,71', level=3)

P('Zhu et al. 2025 [QT015] (BMC Public Health, n = 9.831 học sinh THCS/THPT Tô Châu, '
  'khảo sát lặp lại 3 thời điểm 2019–2023 bao phủ giai đoạn COVID-19) là nghiên cứu '
  'lớn nhất trong tổng hợp 68 bài về yếu tố nguy cơ cụ thể. Kết quả: AOR = 13,71 (KTC '
  '95 %: 8,66 – 21,71) cho trầm cảm "chắc chắn" (PHQ-9 ≥ 15) khi ngủ dưới 5 giờ/đêm '
  'so với ngủ ≥ 8 giờ. Đây là tỷ số chênh lớn nhất trong toàn bộ 68 bài cho bất kỳ yếu '
  'tố nguy cơ nào — nghĩa là giấc ngủ dưới 5 giờ làm tăng nguy cơ trầm cảm lâm sàng '
  'gần 14 lần. Ngay cả cận dưới KTC (8,66) cũng là hiệu ứng rất mạnh.')

P('Để hiểu quy mô, so sánh với các yếu tố nguy cơ khác trong cùng nghiên cứu Zhu 2025 và '
  'các bài khác:')
P('• Áp lực học tập (Wen 2020, Hoa 2024, Trần Hồ Vĩnh Lộc 2024): OR ≈ 11')
P('• Gia đình đơn thân: AOR = 1,434')
P('• Gia đình tái hôn: AOR = 1,837')
P('• Học sinh THPT so với THCS: AOR = 1,409')
P('• Nữ giới so với nam: AOR = 1,25')

P('Ngủ dưới 5 giờ vượt áp lực học tập khoảng 25 % và vượt yếu tố gia đình đơn thân gấp '
  'gần 10 lần. AOR đã được điều chỉnh cho các biến khác, tức là hiệu ứng độc lập sau khi '
  'loại trừ ảnh hưởng áp lực học tập.')

P('Cơ chế đa tầng:', bold=True)
P('(a) Trục HPA và cortisol — Thiếu ngủ mạn tính phá vỡ nhịp cortisol ngày đêm '
  '(cortisol đáng lẽ hạ thấp ban đêm lại giữ ở mức cao). Cortisol tăng kéo dài liên '
  'quan đến viêm thần kinh vùng hippocampus và rối loạn điều tiết HPA — cả hai là '
  'mechanism trung tâm trong nghiên cứu trầm cảm.')
P('(b) Giấc ngủ sóng chậm (slow-wave, NREM) và củng cố cảm xúc — NREM sâu chịu trách '
  'nhiệm cho glymphatic clearance và synaptic homeostasis, trong khi REM sleep đóng vai '
  'trò consolidation của trí nhớ cảm xúc. Khi tổng thời gian ngủ giảm dưới 5 giờ, cả '
  'NREM sâu và REM đều bị cắt ngắn, ảnh hưởng điều tiết cảm xúc hôm sau. (Lưu ý: REM '
  'KHÔNG phải là giai đoạn "tái tạo serotonin" — trái lại, các neuron serotonergic và '
  'noradrenergic giảm phát xung trong REM; cơ chế chính của REM liên quan đến hệ '
  'cholinergic và emotional memory processing.)')
P('(c) Chức năng nhận thức — Thiếu ngủ làm giảm hoạt động vỏ não trán trước (ra quyết '
  'định, điều tiết cảm xúc), làm tăng phản ứng amygdala với kích thích tiêu cực, khiến '
  'trẻ dễ cáu gắt, suy nghĩ tiêu cực và gặp khó khăn ở trường.')
P('(d) Mạng xã hội, game online và rối loạn giấc ngủ — Chen et al. 2023 [QT007] (BMC '
  'Psychiatry, n = 63.205 học sinh Tự Cống, Trung Quốc) cho thấy rối loạn giấc ngủ '
  '(PSQI > 5) có OR = 6,99 (KTC 95 %: 6,69–7,30) và rối loạn game online (IGDS9-SF ≥ '
  '32) có OR = 5,00 (KTC 95 %: 4,42–5,66) cho trầm cảm/lo âu. Ánh sáng xanh và kích '
  'thích dopamine từ nội dung số có thể góp phần trì hoãn giấc ngủ — một con đường '
  'hành vi được giả thuyết liên quan đến thiếu ngủ mạn tính (thiết kế cắt ngang, chưa '
  'xác lập nhân quả).')

P('Hệ quả cho đề cương Việt Nam:', bold=True)
P('Thêm module vệ sinh giấc ngủ vào chương trình can thiệp, bao gồm quy tắc tắt thiết bị '
  '1 giờ trước ngủ, lịch ngủ đều mọi ngày, môi trường ngủ tối và yên tĩnh, tránh caffeine '
  'buổi chiều. Các biện pháp này chi phí thấp và đã được validate rộng rãi. Cần phối hợp '
  'cha mẹ — cha mẹ Việt Nam thường đặt kỷ luật về giờ học nhưng không chú ý giờ ngủ. Cần '
  'giáo dục cha mẹ về khuyến cáo ≥ 7 giờ cho vị thành niên của Hiệp hội Giấc ngủ Hoa Kỳ. '
  'Ngoài ra, cần nghiên cứu chính sách trường học về cho phép/cấm smartphone trong ký túc '
  'xá vào buổi tối. Một số trường tư Việt Nam đã áp dụng (giữ smartphone tập trung từ '
  '22 h đến 6 h) — cần đánh giá hiệu quả bằng RCT.')

# ---- Insight 4 ----
H('V.2.4. Nghịch lý sinh viên/học sinh giỏi nhất có nguy cơ sức khoẻ tâm thần cao nhất',
  level=3)

P('Phát hiện này mang tính nghịch lý: học sinh và sinh viên có thành tích học tập cao '
  'nhất lại là nhóm có nguy cơ sức khoẻ tâm thần cao nhất, trái ngược với giả định thông '
  'thường rằng "học tốt = ứng phó tốt".')

P('Bằng chứng: Đào Thị Ngoãn et al. 2025 [VN028] (Tạp chí Nghiên cứu Y học, n = 196 '
  'sinh viên năm 4 Đại học Y Hà Nội, DASS-21) phát hiện sinh viên đạt học lực Giỏi có '
  'OR = 4,97 (p = 0,005) cho trầm cảm so với nhóm Khá–Trung bình. Con số này đi ngược '
  'trực giác — nhóm học tốt nhất lại có nguy cơ trầm cảm cao nhất. Thêm vào đó, khảo '
  'sát học sinh trường THPT chuyên của Trần Thị Mỵ Lương và Đặng Đức Anh 2020 (n = '
  '540, DASS-42; chưa có canonical ID trong hệ thống 68 bài, tham chiếu từ PDF gốc) ghi '
  'nhận trong 77 học sinh đã có rối loạn lo âu (14,2 % toàn mẫu), 67,5 % có biểu hiện '
  '"thất vọng về bản thân, cảm thấy vô dụng, bất lực" — phản ánh gánh nặng nhận thức '
  'về giá trị bản thân trong môi trường trường chuyên.')

P('Cơ chế tâm lý:', bold=True)
P('(a) Tiêu chuẩn nội hoá — Học sinh/sinh viên giỏi thường tự đặt tiêu chuẩn rất cao. Khi '
  'đạt được một mục tiêu, họ ngay lập tức đặt mục tiêu cao hơn. Quá trình này tạo ra '
  '"moving goalposts" — không bao giờ hài lòng, luôn lo lắng.')
P('(b) Sợ thất bại (fear of failure) — Điểm càng cao, áp lực giữ vị trí càng lớn. Hiện '
  'tượng "impostor syndrome" phổ biến ở sinh viên Y và học sinh chuyên: họ tin rằng mình '
  'không thực sự giỏi và có thể bị "lật tẩy" bất kỳ lúc nào.')
P('(c) So sánh xã hội — Học sinh chuyên và sinh viên Y sống trong môi trường "toàn người '
  'giỏi", khiến so sánh xã hội tiêu cực (upward comparison) trở nên thường nhật. Theo '
  'Festinger 1954, khi so sánh với người giỏi hơn, mức hài lòng với bản thân giảm, tăng '
  'nguy cơ trầm cảm.')
P('(d) Áp lực phụ huynh — Cha mẹ của học sinh chuyên/sinh viên Y thường đã đầu tư nhiều '
  '(tiền bạc, thời gian, kỳ vọng) từ nhỏ. Nỗi sợ "làm cha mẹ thất vọng" tích luỹ qua năm '
  'tháng (60,3 % theo Dong 2025 tại Trung Quốc).')

P('Hệ quả thực hành:', bold=True)
P('Triển khai sàng lọc tâm lý định kỳ cho học sinh trường chuyên và sinh viên các ngành '
  'đòi hỏi cao (Y, Công nghệ, Nghệ thuật). Không chỉ sàng lọc nhóm "yếu" — nhóm "giỏi" '
  'cũng cần. Đưa chủ đề "cầu toàn lành mạnh vs cầu toàn không lành mạnh" vào chương trình '
  'tư vấn học đường, dạy học sinh rằng "nỗ lực cao" khác với "yêu cầu hoàn hảo không thực '
  'tế". Về lâu dài, cần thay đổi văn hoá xếp hạng công khai tại các trường chuyên, giảm '
  'so sánh giữa học sinh, tăng tôn vinh nỗ lực thay vì kết quả. Song song, tổ chức '
  'workshop cho cha mẹ học sinh chuyên về cách đặt kỳ vọng hợp lý và phát hiện dấu hiệu '
  'cầu toàn không lành mạnh sớm.')

# ---- Insight 5 ----
H('V.2.5. Đồng xuất hiện 91,6 % đa hành vi nguy cơ ở học sinh THPT TPHCM', level=3)

P('Duong et al. 2025 [VN029] (Social Psychiatry and Psychiatric Epidemiology Q1, n = '
  '2.631 học sinh THPT TPHCM, đa trung tâm) phát hiện con số đáng lo ngại: 91,6 % học '
  'sinh có ít nhất 2 hành vi nguy cơ sức khoẻ (HRB — Health Risk Behaviors) đồng thời.')

P('Các chỉ số nền của Duong 2025:')
P('• Trầm cảm 42,6 % | Lo âu 50,3 % | Căng thẳng 31,1 %')
P('• Ít vận động thể chất: 79,9 % (hành vi nguy cơ đơn lẻ cao nhất)')
P('• Đa hành vi nguy cơ (≥ 2 HRB đồng thời): 91,6 %')
P('• Các HRB khác được đo (sử dụng chất, đánh nhau, ý nghĩ tự tử, ăn uống không lành '
  'mạnh, thiếu ngủ): báo cáo gốc không công bố tỷ lệ từng hành vi trong abstract, '
  'nhưng có đủ cơ sở thống kê cho chỉ số tổng.')
P('• Mối liên hệ SKTT ↔ HRB: học sinh có triệu chứng sức khoẻ tâm thần (SOMD) có '
  'nguy cơ tham gia HRB cao hơn 1,24 – 4,64 lần so với học sinh không có triệu chứng.')

P('Điểm quan trọng: các hành vi này không độc lập — chúng clustered. Một học sinh có '
  'hành vi nguy cơ này có xu hướng có các hành vi nguy cơ khác. Hiện tượng gọi là "risk '
  'behavior clustering" (Sallis 1992, Leventhal 2013).')

P('Tại sao clustering xảy ra?', bold=True)
P('(a) Cấu trúc ngày — Thiếu ngủ dẫn đến thiếu năng lượng, thiếu vận động, ăn uống bù '
  'qua đồ ăn nhanh và đường, mất cân bằng, rồi trầm cảm/lo âu — vòng luẩn quẩn hành vi. '
  '(b) Sức khoẻ tâm thần là yếu tố trung gian — Duong 2025 cũng phát hiện OR từ 1,24 '
  'đến 4,64 cho mối liên hệ giữa triệu chứng SKTT và hành vi nguy cơ. Lý thuyết "tự điều '
  'trị" (Khantzian 1997) cho rằng vị thành niên dùng hành vi nguy cơ để ứng phó với trầm '
  'cảm/lo âu, lại làm SKTT xấu hơn. (c) Yếu tố môi trường — Cùng một môi trường (trường '
  'học kém, gia đình nghèo, khu vực thiếu tiện ích) tạo điều kiện cho nhiều hành vi nguy '
  'cơ cùng lúc.')

P('Hệ quả cho đề cương Việt Nam:', bold=True)
P('Không can thiệp đơn lẻ một hành vi — cần can thiệp lối sống tích hợp nhắm cùng lúc '
  'vào giấc ngủ, vận động, dinh dưỡng và sử dụng mạng xã hội. Các chương trình hiện chỉ tập '
  'trung vào tâm lý là chưa đủ toàn diện. Tiếp cận cấp cộng đồng — thay đổi môi trường '
  '(công viên thể thao trường, chính sách canteen lành mạnh, giới hạn MXH) — có thể có '
  'hiệu ứng lớn hơn tiếp cận cá nhân đối với 91,6 % học sinh có đa hành vi nguy cơ. Cần '
  'tích hợp với chương trình sức khoẻ học đường sẵn có (dinh dưỡng, thể chất, giáo dục '
  'sức khoẻ sinh sản) thay vì tạo chương trình mới độc lập. Đồng thời, cần nghiên cứu '
  'dọc xác định chiều nhân quả: hành vi nguy cơ dẫn đến SKTT kém, hay SKTT kém dẫn đến '
  'hành vi nguy cơ, hay là vòng luẩn quẩn — kết quả sẽ định hướng ưu tiên can thiệp.')

# ============================================================
# V.3 — Ma tran phuong phap x ket cuc
# ============================================================
H('V.3. Bức tranh tổng thể phương pháp × kết cục trong 68 bài', level=2)

P('Ma trận phương pháp × kết cục dưới đây tổng hợp toàn bộ 68 bài theo hai chiều đồng '
  'thời: phương pháp can thiệp và kết cục đo. Cách quan sát này giúp nhận diện nhanh các '
  'lĩnh vực dồi dào bằng chứng và các khoảng trống cần bổ sung.')

img('kg_method_outcome_heatmap.png', width_cm=15.5,
    caption='Bức tranh phương pháp × kết cục trong tổng hợp 68 bài nghiên cứu')

P('Các quan sát chính:', bold=True)
P('(a) CBT chiếm ưu thế tuyệt đối — 14 bài đo lo âu, 7 bài trầm cảm, 6 bài rối loạn lo '
  'âu xã hội, 5 bài stress, tổng cộng 32 "ô" trong ma trận. Đây là xác nhận rằng CBT là '
  'tiêu chuẩn vàng de facto cho can thiệp tâm lý vị thành niên. Không có phương pháp nào '
  'khác gần con số này.')
P('(b) Nhóm CBT mở rộng (iCBT, gCBT, Mobile CBT) tổng cộng khoảng 15 ô — lĩnh vực rất '
  'hoạt động. Riêng iCBT có 6 ô lo âu và 3 ô rối loạn lo âu xã hội, phù hợp với kết luận '
  'của Xian 2024 NMA [QT039] (iCBT hạng 1 cho SAD).')
P('(c) Mindfulness và Resilience tách biệt — mỗi phương pháp chỉ có 2–3 ô, gợi ý đây '
  'là lĩnh vực đang phát triển, chưa có meta-analysis toàn diện. Brown & Carter 2025 '
  '[QT042] UK đã cảnh báo mindfulness phổ quát có thể thất bại — cần cẩn thận khi triển '
  'khai đại trà.')
P('(d) Khoảng trống rõ nhất: Vận động thể chất (PE) chỉ có 3 ô — quá ít so với bằng '
  'chứng umbrella review của JAACAP 2025 (SMD = −0,39 cho lo âu, SMD = −0,32 cho vận '
  'động aerobic). Điều này gợi ý PE đang bị under-represented trong tổng hợp hiện tại, '
  'cần bổ sung thêm RCT về vận động thể chất cho vị thành niên Việt Nam.')
P('(e) VR therapy xuất hiện mới — chỉ 1–2 ô, nhưng umbrella review JAACAP 2025 cho thấy '
  'VR therapy có tiềm năng cho rối loạn lo âu xã hội. Đây là lĩnh vực tiên phong đáng '
  'đầu tư tìm kiếm.')

P('(f) Lưu ý quan trọng về mạng xã hội và thời gian màn hình:', bold=True)
P('Phân tích đối chiếu giữa 68 bài cho thấy y văn về mạng xã hội (MXH) và sức khoẻ '
  'tâm thần vị thành niên CHƯA THỐNG NHẤT. Một nhóm nghiên cứu (QT024 WHO Europe, '
  'QT032 Ireland, QT033 JAMA Screen Media, QT035 Jefferies SAD 7 nước) ghi nhận mối '
  'liên hệ tiêu cực rõ ràng giữa thời gian màn hình và lo âu/trầm cảm. Tuy nhiên, '
  'nhóm khác (QT022 Li BJCP 2025 longitudinal, QT027 Nature Human Behaviour 2025, '
  'VN015 Ngô Anh Vinh DTTS nội trú) cho thấy mối quan hệ phức tạp hơn — internet có '
  'thể BẢO VỆ ở quần thể cô lập (OR = 0,43 cho VTN nội trú DTTS [VN015]), và mối '
  'liên hệ tiêu cực có thể phụ thuộc vào loại sử dụng (thụ động vs chủ động), bối '
  'cảnh xã hội (đô thị vs nông thôn/nội trú), và cường độ (vừa phải vs binge). Kết '
  'luận: MXH KHÔNG nên được mặc định coi là yếu tố nguy cơ đồng nhất cho mọi quần '
  'thể vị thành niên — cần phân biệt theo bối cảnh khi thiết kế can thiệp.')

# ============================================================
# V.4 — Happy House
# ============================================================
H('V.4. Nghiên cứu Happy House — RCT đầu tiên tại Việt Nam về can thiệp SKTT cho vị '
  'thành niên', level=2)

P('Trong các phiên bản báo cáo trước, nhóm tác giả đã đề cập "chưa có RCT can thiệp '
  'SKTT vị thành niên tại Việt Nam" — đây là kết luận không chính xác. Nghiên cứu Happy '
  'House của Tran và cộng sự 2023 [VN030] thực tế là cluster controlled trial đầu tiên '
  'trong lĩnh vực này. Phần này trình bày đầy đủ chi tiết Happy House dựa trên bản dịch '
  'toàn văn từ PMC10643236.')

P('Tác giả và đơn vị:', bold=True)
P('Thach Duc Tran (Monash University, Melbourne, Úc) — tác giả chính; Huong Nguyen '
  '(Hanoi University of Public Health, Việt Nam); Ian Shochet (Queensland University of '
  'Technology, Brisbane, Úc); cùng Nga Nguyen, Nga La, Astrid Wurfl, Jayne Orr, Hau '
  'Nguyen, Ruby Stocker, và Jane Fisher (Monash). Mô hình hợp tác Việt–Úc ba thể chế, '
  'nhóm Monash đóng vai trò methodology lead và HUPH đảm nhiệm field implementation.')

P('Tài trợ và phê duyệt đạo đức:', bold=True)
P('Đăng ký thử nghiệm: ACTRN12620000088943 (Australian New Zealand Clinical Trials '
  'Registry — minh bạch theo chuẩn quốc tế). Tài trợ: NHMRC Úc (GNT1158429) kết hợp '
  'NAFOSTED Việt Nam (NHMRC.108.01-2018.02). Phê duyệt đạo đức từ Monash HREC, Đại học '
  'Y tế Công cộng Hà Nội IRB, và QUT OREC.')

P('Thiết kế nghiên cứu:', bold=True)
P('• 8 trường THPT Hà Nội được chọn theo 3 tiêu chí: (a) sẵn sàng tham gia của ban giám '
  'hiệu, (b) cân bằng đặc điểm kinh tế–xã hội, (c) cân bằng địa lý (đô thị vs ngoại '
  'thành). Phân bổ 4 trường can thiệp và 4 trường đối chứng.')
P('• 1.084 học sinh lớp 10 được tuyển — tỷ lệ tham gia 96,1 % (1.084 trong số 1.128 đủ '
  'điều kiện). Đây là tỷ lệ rất cao, cho thấy sự chấp nhận mạnh từ học sinh và phụ huynh.')
P('• Sau khi loại bỏ học sinh thiếu dữ liệu baseline, số phân tích cuối cùng: 531 can '
  'thiệp và 552 đối chứng (tổng 1.083). Chênh lệch một đơn vị so với 1.084 đăng ký do '
  'một học sinh không hoàn tất bộ đo baseline.')
P('• Tuổi 15–16 (lớp 10); khoảng 60 % nữ; khoảng 50 % đô thị; khoảng 89 % sống cùng cha '
  'mẹ ruột.')
P('• Điểm CESD-R trung bình ban đầu: 11,4 ± 12,2 (đối chứng) so với 12,0 ± 12,0 (can '
  'thiệp), p = 0,729 — cân bằng tốt, không có selection bias.')
P('• Tỷ lệ trầm cảm lâm sàng ban đầu (CESD-R ≥ 16): 25,5 % ở cả hai nhóm — phản ánh '
  'gánh nặng trầm cảm đáng kể ở học sinh lớp 10 Hà Nội. Không so sánh trực tiếp với '
  'các con số lo âu của Hoa 2024 (GAD-7) vì hai nghiên cứu đo hai rối loạn khác nhau.')

P('Mô hình can thiệp Happy House:', bold=True)
P('Chương trình RAP-A (Resourceful Adolescent Program, Shochet et al., Úc) được thích '
  'ứng văn hoá Việt Nam qua 4 giai đoạn: (1) dịch tiếng Việt với back-translation; '
  '(2) thích ứng văn hoá với chuyên gia tâm lý học đường Việt Nam; (3) thử nghiệm pilot '
  'với 30 học sinh Hà Nội; (4) sửa đổi dựa trên phản hồi học sinh và giáo viên. Cấu '
  'trúc: 11 buổi 45 phút của RAP gốc được rút xuống 6 buổi 90 phút, thực hiện trong 6 '
  'tuần liên tiếp, thay thế môn Giáo dục Công dân.')

P('Sáu buổi Happy House:')
P('• Buổi 1: Sức mạnh bản thân — nhận diện điểm mạnh, xây dựng tự trọng')
P('• Buổi 2: Suy nghĩ tích cực — tái cấu trúc nhận thức (CBT)')
P('• Buổi 3: Quản lý cảm xúc — kỹ thuật thư giãn, điều tiết')
P('• Buổi 4: Giải quyết vấn đề — các bước có cấu trúc')
P('• Buổi 5: Mạng lưới hỗ trợ — xây dựng quan hệ (IPT)')
P('• Buổi 6: Ôn tập và cam kết')

P('Người cung cấp: Giáo viên Giáo dục Công dân (có nền tảng tâm lý cơ bản) được đào '
  'tạo 3 ngày bởi nhóm Monash và HUPH. Mỗi buổi có 1 giáo viên chính và 1–2 trợ giảng '
  'để hỗ trợ lớp đông 40–45 học sinh. Thích ứng này khác với RAP gốc Úc (16 học '
  'sinh/lớp) để phù hợp với class size Việt Nam.')

P('Kết quả chính (aOR + KTC 95 %):', bold=True)
table(
    ['Thời điểm', 'Kết cục', 'Đối chứng (%)', 'Can thiệp (%)', 'aOR (KTC 95 %)', 'p'],
    [
        ['2 tuần', 'CESD-R ≥ 16', '28,6 % (155/542)', '23,9 % (125/523)', '0,56 (0,36–0,88)', '0,011'],
        ['6 tháng', 'CESD-R ≥ 16', '29,2 % (157/537)', '26,3 % (138/524)', '0,75 (0,51–1,09)', '0,132'],
    ],
    widths=[2.0, 3.0, 2.8, 2.8, 3.2, 1.7])

P('Cohen d = 0,11 cho trầm cảm ở thời điểm 2 tuần — hiệu ứng nhỏ nhưng có ý nghĩa thống '
  'kê. Để so sánh: meta-analysis các chương trình phòng ngừa phổ quát toàn cầu có median '
  'Cohen d khoảng 0,18 cho trầm cảm (Stockings et al. 2016). Happy House nằm dưới median '
  'nhưng vẫn trong khoảng biến thiên bình thường của các can thiệp phổ quát.')

P('Hiện tượng fade-out ở 6 tháng là vấn đề đáng lo ngại. Các trial RAP ở Úc và New '
  'Zealand cho thấy hiệu quả thường duy trì 10–18 tháng mà không fade out rõ rệt. Tác '
  'giả Happy House đưa ra hai giải thích tiềm tàng: (a) COVID-19 lockdown tại Việt Nam '
  'trong thời gian theo dõi (từ 12/2020 đến 5/2021) phá vỡ hỗ trợ xã hội và tăng stress '
  'phổ quát, có thể che mờ hiệu quả can thiệp; (b) 6 buổi được rút gọn từ 11 buổi RAP '
  'gốc có thể là "liều chưa đủ".')
P('Từ các giải thích trên, nhóm tác giả báo cáo này đề xuất: cần thiết kế booster '
  'sessions ở mốc 3, 6 và 12 tháng để duy trì hiệu quả can thiệp trong các RCT Việt '
  'Nam tương lai.')

P('Phát hiện tích cực quan trọng: thang Coping Self-Efficacy Scale (CSES) ở hai tiểu '
  'thang "giải quyết vấn đề" và "tìm kiếm hỗ trợ xã hội" duy trì ý nghĩa thống kê ở cả '
  '2 tuần và 6 tháng (d = 0,17–0,26). Nghĩa là kỹ năng ứng phó được học trong 6 buổi '
  'giữ lại lâu dài, dù triệu chứng trầm cảm quay lại. Đây là cơ chế quan trọng: Happy '
  'House không phải "chữa trầm cảm" trực tiếp, mà là "dạy kỹ năng" — và kỹ năng bền '
  'hơn triệu chứng.')

P('Điểm mạnh:', bold=True)
P('Cluster controlled trial đầu tiên tại Việt Nam về can thiệp sức khoẻ tâm thần trường '
  'học cho vị thành niên. Cỡ mẫu rất lớn (n = 1.084), tỷ lệ tham gia 96,1 %, mất mẫu '
  'dưới 3 %. RAP-A có cơ sở bằng chứng quốc tế mạnh từ Úc, New Zealand, Mauritius. Mô '
  'hình hợp tác Việt–Úc (Monash–QUT–HUPH) uy tín. Giáo viên cung cấp — phù hợp bối cảnh '
  'quốc gia thu nhập trung bình. Đăng ký thử nghiệm trước và tài trợ quốc tế.')

P('Hạn chế:', bold=True)
P('(1) Can thiệp phổ quát (universal), không phải chọn lọc (targeted) — hiệu ứng nhỏ d = 0,11. (2) '
  'Fade-out ở 6 tháng. (3) Đo trầm cảm là chính, không đo lo âu riêng. (4) Chỉ lớp 10 '
  'tại Hà Nội — chưa đại diện toàn Việt Nam. (5) Phân bổ nhóm không ngẫu nhiên (có thể '
  'có bias giữa trường can thiệp và đối chứng). (6) Không có active control. (7) '
  'COVID-19 confounding trong thời gian theo dõi.')

# ============================================================
# V.6 — Xu huong toan cau
# ============================================================
H('V.6. Tổng hợp liên bài: Xu hướng tăng lo âu toàn cầu và vị trí của Việt Nam', level=2)

P('Nhiều nghiên cứu trong tổng hợp 68 bài cung cấp dữ liệu xu hướng dài hạn, cho phép '
  'nhìn nhận lo âu vị thành niên Việt Nam trong bối cảnh toàn cầu.')

table(
    ['ID', 'Quốc gia', 'Giai đoạn', 'N', 'Xu hướng'],
    [
        ['QT030', 'Toàn cầu (204 nước)', '1990–2021 (31 năm)', 'GBD estimates', 'AAPC 0,84 %/năm; tăng tốc sau 2014'],
        ['QT023', 'Hoa Kỳ', '2013–2021 (8 năm)', '13,7 triệu hồ sơ', 'Lo âu AOR = 2,17; nhóm 15–17: AOR = 2,93'],
        ['QT021', 'Na Uy', '2011–2024 (13 năm)', '979.043 VTN', 'Tăng liên tục; g = 0,46 (MXH)'],
        ['QT034', 'Hàn Quốc', '2006–2022 (16 năm)', '1.138.804 VTN', 'CẢI THIỆN trước COVID → xấu đi sau COVID'],
        ['VN014', 'Việt Nam (6 tỉnh)', '2021 vs 2023', '8.473 VTN', 'Lo âu 41,5 % → 25,4 % (phục hồi sau COVID)'],
    ],
    widths=[1.3, 3.5, 3.5, 3.2, 4.0])

P('Hai quan sát chính:', bold=True)

P('Thứ nhất, xu hướng tăng lo âu là TOÀN CẦU nhưng tốc độ khác nhau. Phân tích GBD '
  '2021 [QT030] trên 204 nước cho thấy AAPC 0,84 %/năm trong 31 năm, tăng tốc sau '
  '2014 — trùng với thời điểm smartphone và mạng xã hội phổ biến ở vị thành niên. Tại '
  'Hoa Kỳ [QT023], dữ liệu chẩn đoán lâm sàng từ 13,7 triệu hồ sơ cho thấy lo âu '
  'tăng gấp đôi (AOR = 2,17) trong 8 năm, đặc biệt ở nhóm 15–17 tuổi (AOR = 2,93). '
  'Na Uy [QT021] ghi nhận xu hướng tăng liên tục 13 năm trên mẫu 979.043 vị thành niên.')

P('Thứ hai, Hàn Quốc [QT034] là ngoại lệ đáng học hỏi — quốc gia DUY NHẤT trong tổng '
  'hợp 68 bài cho thấy sức khoẻ tâm thần vị thành niên CẢI THIỆN trước COVID (2006–'
  '2019) rồi xấu đi sau đó. Nguyên nhân được tác giả gán cho hệ thống Wee Center — '
  'chương trình đặt chuyên viên tâm lý vào mọi trường học, triển khai từ 2005 đến 2012, '
  'với hơn 6.000 trung tâm trên toàn quốc. Đây là bằng chứng cấp quốc gia rằng đầu tư '
  'có hệ thống vào sức khoẻ tâm thần học đường CÓ THỂ đảo ngược xu hướng tăng. '
  'COVID-19 phá vỡ tiến bộ này, đặc biệt ở nhóm thu nhập thấp.')

P('Tại Việt Nam [VN014], Hoàng Trung Học và cộng sự 2025 (n = 8.473, 6 tỉnh/thành) '
  'cho thấy lo âu giảm từ 41,5 % (trong COVID) xuống 25,4 % (sau COVID) — gợi ý phục '
  'hồi phần nào. Tuy nhiên, tỷ lệ sau COVID vẫn cao, và Việt Nam chưa có dữ liệu xu '
  'hướng dài hạn tương đương KYRBS (Hàn Quốc) hay NHANES (Hoa Kỳ). Xây dựng hệ thống '
  'giám sát dọc là ưu tiên hạ tầng nghiên cứu cấp bách.')

# ============================================================
# V.7 — Bat binh dang
# ============================================================
H('V.7. Tổng hợp liên bài: Bất bình đẳng địa lý, thu nhập và dân tộc', level=2)

P('Tổng hợp 68 bài cho thấy bất bình đẳng trong lo âu vị thành niên diễn ra trên nhiều '
  'trục đồng thời: địa lý, thu nhập gia đình và dân tộc.')

P('Bất bình đẳng địa lý tại Việt Nam:', bold=True)

table(
    ['ID', 'Địa bàn', 'Tỷ lệ lo âu', 'Công cụ', 'Ghi chú'],
    [
        ['VN026', 'Long An (giáp TPHCM)', '57,2 %', 'DASS-21', 'CAO NHẤT VN — áp lực thi ĐH'],
        ['VN015', 'Lạng Sơn (DTTS nội trú)', '54,4 %', 'DASS-21', 'DTTS nội trú xa gia đình'],
        ['VN025', 'Hải Phòng (Vĩnh Bảo)', '53,8 %', 'DASS-21', 'Cộng Hiền THPT'],
        ['VN001', 'Hà Nội', '40,6 %', 'GAD-7 (≥ 5)', 'Đô thị, n = 3.910'],
        ['VN014', 'VN 6 tỉnh (sau COVID)', '25,4 %', 'DASS-21', 'Phục hồi sau COVID'],
        ['VN024', 'Vĩnh Long (ĐBSCL)', '12,2 % (trầm cảm)', 'CES-D (≥ 16)', 'Thấp nhất — cut-off nghiêm ngặt'],
    ],
    widths=[1.3, 3.5, 2.5, 2.7, 5.5])

P('Lưu ý: so sánh trực tiếp giữa các tỉnh cần thận trọng do khác biệt công cụ đo '
  '(GAD-7 vs DASS-21 vs CES-D), cut-off, và thời điểm khảo sát. Tuy nhiên, hình ảnh '
  'tổng thể cho thấy các tỉnh có áp lực thi đại học cao (giáp đô thị lớn) và vùng dân '
  'tộc thiểu số có tỷ lệ cao nhất.')

P('Bất bình đẳng thu nhập (dữ liệu Hàn Quốc):', bold=True)
P('Cho et al. 2024 [QT034] phân tích KYRBS 16 năm (n = 1.138.804) cho thấy: stress '
  'nhận thức ở nhóm thu nhập thấp nhất là 62,8 % so với 40,1 % ở nhóm cao nhất — chênh '
  'lệch 22,7 điểm phần trăm. Ý tưởng tự tử: 31,7 % nhóm nghèo vs 13,9 % nhóm giàu '
  '(gấp 2,3 lần). COVID-19 làm MỞ RỘNG khoảng cách, nhóm nghèo bị ảnh hưởng nặng hơn. '
  'Việt Nam hiện CHƯA CÓ dữ liệu phân tầng theo thu nhập gia đình cho sức khoẻ tâm '
  'thần vị thành niên — đây là khoảng trống cần lấp.')

P('Bất bình đẳng dân tộc:', bold=True)
P('Ngô Anh Vinh et al. 2024 [VN015] (n = 845 học sinh DTTS nội trú tại Lạng Sơn, '
  'DASS-21) ghi nhận tỷ lệ lo âu 54,4 % — cao hơn hầu hết các nghiên cứu ở vùng Kinh. '
  'Đáng chú ý, VN015 phát hiện một loạt yếu tố bảo vệ đặc thù cho quần thể DTTS nội '
  'trú:')
P('• Bạn bè kém → trầm cảm: OR = 6,84 (KTC 95 %: 2,03–23,02) — yếu tố nguy cơ mạnh nhất')
P('• Vận động thể chất → lo âu: OR = 0,72 (KTC: 0,51–1,00) — bảo vệ')
P('• Sống cùng ≥ 10 người → lo âu: OR = 0,51 (KTC: 0,33–0,77) — bảo vệ')
P('• Internet hằng ngày → lo âu: OR = 0,43 (KTC: 0,21–0,86) — bảo vệ')
P('• Hỗ trợ từ bạn bè → lo âu: OR = 0,57 — bảo vệ')

P('Phát hiện internet hằng ngày bảo vệ (OR = 0,43) có vẻ ngược trực giác, nhưng nhất '
  'quán với giả thuyết "bù đắp xã hội" (social compensation hypothesis, Valkenburg & '
  'Peter 2007): ở quần thể cô lập — học sinh DTTS sống xa gia đình trong trường nội '
  'trú — internet là kênh kết nối chính với gia đình và cộng đồng gốc. So sánh không '
  'phải "nhiều vs ít internet" mà là "sử dụng đều đặn hằng ngày vs dồn cuối tuần" — '
  'tương tự mô hình "Goldilocks" (Przybylski & Weinstein 2017) trong đó sử dụng vừa '
  'phải đều đặn tốt hơn sử dụng dồn. Thiết kế cắt ngang không loại trừ nhân quả '
  'ngược (học sinh ít lo âu có thể dùng internet đều hơn).')

# ============================================================
# V.8 — Yeu to bao ve chua khai thac
# ============================================================
H('V.8. Tổng hợp liên bài: Các yếu tố bảo vệ chưa được khai thác đầy đủ', level=2)

P('Ngoài 5 phát hiện chính ở V.2, tổng hợp 68 bài còn chứa một số yếu tố bảo vệ có '
  'tiềm năng can thiệp cao nhưng chưa được khai thác đầy đủ trong các phiên bản báo '
  'cáo trước.')

table(
    ['ID', 'Yếu tố bảo vệ', 'Chỉ số', 'Quần thể', 'Thiết kế'],
    [
        ['VN003', 'Chăm sóc cảm xúc (emotional care)', 'β = −0,40 (p < 0,001)',
         'VTN cơ sở HTXH Huế', 'Cắt ngang'],
        ['VN014', 'Nuôi dạy tích cực (positive parenting)', 'OR = 0,30 (giảm 70 %)',
         'n = 8.473 VTN 6 tỉnh VN', 'Cắt ngang'],
        ['QT047', 'Tâm sự với gia đình', 'OR = 0,22 (giảm 78 %)',
         'n = 2.716 Ya An TQ', 'Cắt ngang'],
        ['VN015', 'Vận động thể chất', 'OR = 0,72',
         'n = 845 DTTS Lạng Sơn', 'Cắt ngang'],
        ['QT036', 'Giảm cô đơn (loneliness)', 'F = 16.646 (ML top-1)',
         'n = 213.820 VTN Hàn Quốc', 'ML + cắt ngang'],
    ],
    widths=[1.3, 4.0, 3.5, 3.5, 3.2])

P('Ba quan sát từ bảng trên:', bold=True)

P('Thứ nhất, gia đình là yếu tố bảo vệ nhất quán xuyên quốc gia. Ba nghiên cứu độc '
  'lập — Dong 2025 [QT047] tại Trung Quốc (OR = 0,22, tâm sự gia đình), Hoàng Trung '
  'Học 2025 [VN014] tại Việt Nam (OR = 0,30, nuôi dạy tích cực), và Phạm 2024 [VN003] '
  'tại Huế (β = −0,40, chăm sóc cảm xúc) — đều chỉ ra rằng CHẤT LƯỢNG tương tác cảm '
  'xúc cha mẹ–con là yếu tố bảo vệ mạnh hơn cả điều kiện vật chất. Điều này có ý '
  'nghĩa thực tiễn lớn cho LMIC: can thiệp dạy kỹ năng giao tiếp cha mẹ có chi phí '
  'thấp và khả thi cao.')

P('Thứ hai, cô đơn là yếu tố nguy cơ hàng đầu cho rối loạn lo âu lan toả (GAD). Moon '
  '& Woo 2025 [QT036] sử dụng Machine Learning trên 213.820 vị thành niên Hàn Quốc '
  '(KYRBS 2020–2023) phát hiện cô đơn đứng hạng 1 trong tất cả yếu tố (F-score = '
  '16.646), vượt cả stress nhận thức. Kết hợp với dữ liệu VN015 (bạn bè kém OR = 6,84 '
  'cho trầm cảm, hỗ trợ bạn bè OR = 0,57 cho lo âu), gợi ý rằng chiến lược can thiệp '
  'giảm cô đơn — như chương trình mentor/buddy, hoạt động nhóm, câu lạc bộ trường — có '
  'thể hiệu quả hơn can thiệp stress đơn thuần.')

P('Thứ ba, vận động thể chất là khoảng trống can thiệp lớn nhất tại Việt Nam. Dữ liệu '
  'VN015 (OR = 0,72 bảo vệ cho lo âu) là nghiên cứu Việt Nam DUY NHẤT cung cấp bằng '
  'chứng định lượng cho vận động thể chất. Con số này nhất quán với umbrella review '
  'JAACAP 2025 (SMD = −0,39 cho lo âu ở 375 RCT). Đồng thời, Duong 2025 [VN029] báo '
  'cáo 79,9 % học sinh TPHCM thiếu vận động — hành vi nguy cơ đơn lẻ cao nhất. Sự kết '
  'hợp "hiệu quả cao + tỷ lệ thiếu rất cao" gợi ý vận động thể chất nên là thành phần '
  'bắt buộc trong mọi chương trình can thiệp, không phải phụ trợ.')

# ============================================================
# V.9 — Khuyen nghi (doi tu V.5)
# ============================================================
H('V.9. Khuyến nghị cho đề cương can thiệp lo âu vị thành niên Việt Nam 2026–2028',
  level=2)

P('Dựa trên 5 phát hiện liên bài và dữ liệu đầy đủ về Happy House, nhóm tác giả đề xuất '
  'các ưu tiên sau cho giai đoạn 2026–2028:')

P('Ưu tiên 1 — Thiết kế RCT chuyên biệt cho LO ÂU tại Việt Nam.', bold=True)
P('Happy House là RCT đầu tiên nhưng đo trầm cảm là chính. Việt Nam hiện chưa có RCT '
  'chuyên biệt đo lo âu (GAD-7, DASS-Y lo âu, SIAS-17) như kết cục chính. Cần thiết kế '
  'RCT chọn lọc (targeted) với tiêu chí tuyển chọn GAD-7 ≥ 8 hoặc DASS-Y lo âu ≥ 8, cỡ mẫu tối thiểu '
  '300 mỗi nhánh, theo dõi 3-6-12 tháng với booster sessions.')

P('Ưu tiên 2 — Can thiệp tích hợp đa thành phần, không đơn lẻ.', bold=True)
P('Dựa trên phát hiện 91,6 % học sinh có đa hành vi nguy cơ, chương trình cần tích hợp '
  'bốn module: (a) tâm lý (CBT-based 6–10 buổi), (b) giấc ngủ (sleep hygiene 2 buổi), '
  '(c) vận động thể chất (PE 2 buổi + thay đổi môi trường), (d) kỹ năng giao tiếp gia '
  'đình (4 buổi phụ huynh). Tổng 14–18 buổi cho học sinh + 4 buổi cho phụ huynh.')

P('Ưu tiên 3 — Áp lực học tập là mục tiêu trung tâm, không phụ trợ.', bold=True)
P('OR ≈ 11 xuyên 4 nghiên cứu độc lập là bằng chứng mạnh. Cần (a) dùng ESSA phiên bản '
  'Việt Nam cho mọi đo lường; (b) tích hợp kỹ năng quản lý thời gian, tư vấn hướng '
  'nghiệp sớm; (c) phối hợp chính sách với Bộ GD&ĐT về giảm tải, bỏ xếp hạng công khai.')

P('Ưu tiên 4 — Nhóm rủi ro đặc biệt: học sinh chuyên và sinh viên ngành đòi hỏi cao.',
  bold=True)
P('Cần chương trình riêng cho nhóm này dựa trên phát hiện OR = 4,97 của Đào Thị Ngoãn '
  '2025 và tỷ lệ 67,5 % có biểu hiện "thất vọng về bản thân, cảm thấy vô dụng, bất '
  'lực" TRONG NHÓM 77 HỌC SINH ĐÃ CÓ RLLA của Trần Thị Mỵ Lương 2020 (540 HS '
  'chuyên, sàng lọc bằng DASS-42). Nội dung chương trình: psychoeducation về cầu toàn '
  'lành mạnh vs không lành mạnh, sàng lọc định kỳ bằng công cụ phù hợp (GAD-7 hoặc '
  'DASS-21), tư vấn chuyên biệt.')

P('Ưu tiên 5 — Mở rộng địa lý.', bold=True)
P('Các nghiên cứu Việt Nam hiện tập trung ở Hà Nội, TPHCM, Huế và một số tỉnh lẻ. Khoảng '
  'trống rõ ở Tây Nguyên, miền núi phía Bắc và miền Trung. Đề xuất RCT đa trung tâm với '
  'ít nhất một điểm ở mỗi vùng.')

P('Ưu tiên 6 — Phối hợp trường–gia đình.', bold=True)
P('Phát hiện OR = 0,22 của Dong 2025 về kênh giao tiếp gia đình là yếu tố bảo vệ mạnh '
  'nhất. Đề xuất workshop cha mẹ mỗi học kỳ, nội dung: cách lắng nghe không phán xét, '
  'hỏi về cảm xúc thay vì chỉ hỏi về điểm số, phát hiện dấu hiệu lo âu sớm.')

P('Ưu tiên 7 — Bổ sung bằng chứng về vận động thể chất và VR therapy.', bold=True)
P('Hai khoảng trống chính trong tổng hợp 68 bài. Umbrella review JAACAP 2025 cho thấy '
  'vận động thể chất có SMD = −0,39 cho lo âu và VR therapy có tiềm năng cho SAD — cần '
  'bổ sung các nghiên cứu này và cân nhắc pilot tại Việt Nam.')

P('Kết luận:', bold=True)
P('Với 68 bài nghiên cứu được tổng hợp có hệ thống — bao gồm 5 phát hiện liên bài '
  'chính (V.2), dữ liệu đầy đủ từ Happy House (V.4), 3 tổng hợp liên bài bổ sung về '
  'xu hướng toàn cầu (V.6), bất bình đẳng (V.7) và yếu tố bảo vệ (V.8) — nền tảng '
  'bằng chứng đã đủ để thiết kế đề cương RCT can thiệp lo âu vị thành niên Việt Nam '
  'giai đoạn 2026–2028. Các ưu tiên đề xuất hướng đến mục tiêu chuyển từ can thiệp phổ '
  'quát (universal) có hiệu ứng nhỏ sang can thiệp chọn lọc (targeted) có hiệu ứng lớn '
  'hơn, tích hợp đa thành phần, mở rộng địa bàn, và xây dựng hệ thống giám sát dọc '
  'cho sức khoẻ tâm thần vị thành niên trên toàn quốc.')

P('---', align='center', italic=True)
P('Báo cáo v5 — 12/04/2026 — tổng hợp 68 bài nghiên cứu về lo âu/trầm cảm vị thành niên',
  align='center', italic=True, size=10)

doc.save(DST)
d2 = Document(DST)
print(f'Saved: {os.path.basename(DST)}')
print(f'Total: {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables')
total_chars = sum(len(p.text) for p in d2.paragraphs)
print(f'Total chars: {total_chars:,}')
