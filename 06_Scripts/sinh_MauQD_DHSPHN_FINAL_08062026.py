# -*- coding: utf-8 -*-
"""Ban mau QD HD Dao duc FINAL cho Truong DHSPHN — gui CT Hang xuc tien thu tuc.
Framework: phan tich thu cap + cong bo quoc te tu chuong trinh sang loc thuong nien.
3 ten bai chinh thuc da chot ngay 08/06/2026.
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'MauQD_HoiDongDaoDuc_DHSPHN_FINAL_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(1.7); sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.7)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.3

RED = RGBColor(0xC0, 0x00, 0x00)


# ============================================================
# Header (2 cot)
# ============================================================
table = d.add_table(rows=1, cols=2)
table.autofit = False
table.columns[0].width = Cm(7.5); table.columns[1].width = Cm(7.5)
cell1 = table.cell(0, 0); cell2 = table.cell(0, 1)
cell1.paragraphs[0]._element.getparent().remove(cell1.paragraphs[0]._element)
cell2.paragraphs[0]._element.getparent().remove(cell2.paragraphs[0]._element)


def add_centered(cell, text, sz=11, bold=False, italic=False, is_placeholder=False):
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text); r.font.name = 'Times New Roman'
    r.font.size = Pt(sz)
    if bold: r.bold = True
    if italic: r.italic = True
    if is_placeholder:
        r.font.color.rgb = RED
        # Bold inherited from caller — em không force


add_centered(cell1, 'BỘ GIÁO DỤC VÀ ĐÀO TẠO', 11, bold=True)
add_centered(cell1, 'TRƯỜNG ĐẠI HỌC SƯ PHẠM HÀ NỘI', 11, bold=True)
add_centered(cell1, '—————', 11)
add_centered(cell1, 'Hội đồng Đạo đức trong Nghiên cứu Khoa học', 11, bold=True)
# Số QĐ — placeholder
p = cell1.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Số: '); r.font.size = Pt(11); r.bold = True
r2 = p.add_run('[___/QĐ-ĐHSPHN]'); r2.font.size = Pt(11); r2.bold = True
r2.font.color.rgb = RED

add_centered(cell2, 'CỘNG HÒA XÃ HỘI CHỦ NGHĨA VIỆT NAM', 11, bold=True)
add_centered(cell2, 'Độc lập - Tự do - Hạnh phúc', 11, bold=True)
add_centered(cell2, '—————', 11)
# Ngày — placeholder không đậm
p = cell2.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Hà Nội, ngày '); r.font.size = Pt(11); r.italic = True
r2 = p.add_run('[___]'); r2.font.size = Pt(11); r2.italic = True
r2.font.color.rgb = RED
r3 = p.add_run(' tháng '); r3.font.size = Pt(11); r3.italic = True
r4 = p.add_run('[___]'); r4.font.size = Pt(11); r4.italic = True
r4.font.color.rgb = RED
r5 = p.add_run(' năm 2025'); r5.font.size = Pt(11); r5.italic = True


# ============================================================
# Tieu de
# ============================================================
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
r = p.add_run('QUYẾT ĐỊNH'); r.bold = True; r.font.size = Pt(14)

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_after = Pt(6)
r = p.add_run('V/v: Phê duyệt phân tích thứ cấp và công bố quốc tế đối với '
              'dữ liệu sàng lọc sức khỏe tâm thần học sinh trung học cơ sở')
r.italic = True; r.font.size = Pt(11)


# ============================================================
# Chu tich + Can cu
# ============================================================
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
r = p.add_run('CHỦ TỊCH HỘI ĐỒNG ĐẠO ĐỨC TRONG NGHIÊN CỨU KHOA HỌC')
r.bold = True; r.font.size = Pt(12)


def CC(text_segments):
    """Một dòng căn cứ. text_segments là list of (text, is_placeholder)."""
    p = d.add_paragraph()
    p.paragraph_format.first_line_indent = Cm(0.75); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    for text, is_ph in text_segments:
        r = p.add_run(text)
        r.font.size = Pt(11); r.italic = True
        if is_ph:
            r.font.color.rgb = RED


CC([
    ('Căn cứ Thông tư số 43/2024/TT-BYT ngày 12/12/2024 của Bộ trưởng '
     'Bộ Y tế quy định việc thành lập, tổ chức và hoạt động của Hội '
     'đồng Đạo đức trong nghiên cứu y sinh học (áp dụng tham chiếu cho '
     'nghiên cứu khoa học giáo dục/tâm lý học);', False),
])
CC([
    ('Căn cứ Quyết định số ', False),
    ('[___/QĐ-ĐHSPHN]', True),
    (' ngày ', False),
    ('[___]', True),
    (' của Hiệu trưởng Trường Đại học Sư phạm Hà Nội về việc thành '
     'lập Hội đồng Đạo đức trong Nghiên cứu Khoa học của Trường;', False),
])
CC([
    ('Căn cứ Hồ sơ xin phê duyệt đạo đức đề tài nghiên cứu của nghiên '
     'cứu sinh Công Thị Hằng đề ngày ', False),
    ('[___]', True),
    (' tháng ', False),
    ('[___]', True),
    (' năm 2025, kèm theo Thư đồng ý cho phép thu thập dữ liệu trong '
     'khuôn khổ chương trình sàng lọc sức khỏe tâm thần học sinh '
     'thường niên của Ban Giám hiệu hai Trường Trung học cơ sở Nhật '
     'Tân và Tây Mỗ;', False),
])
CC([
    ('Căn cứ kết luận của Hội đồng Đạo đức tại Phiên họp thẩm định '
     'ngày ', False),
    ('[___/___/2025]', True),
    (';', False),
])
CC([
    ('Theo đề nghị của Thư ký Hội đồng Đạo đức,', False),
])

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
r = p.add_run('QUYẾT ĐỊNH:'); r.bold = True; r.font.size = Pt(13)


# ============================================================
# Cac dieu
# ============================================================
def DIEU(num, segments):
    """Một điều với label đậm + nội dung."""
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(3)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(f'Điều {num}. ')
    r.bold = True; r.font.size = Pt(11)
    for text, is_ph in segments:
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        if is_ph:
            r2.font.color.rgb = RED


def SUB(label, value_segments):
    """Bullet • Label: value với value có thể chứa placeholder."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(f'• {label}: '); r.bold = True; r.font.size = Pt(11)
    for text, is_ph in value_segments:
        r2 = p.add_run(text)
        r2.font.size = Pt(11)
        if is_ph:
            r2.font.color.rgb = RED


def BULL(text, is_ph=False, last=False):
    """Bullet đơn."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0); p.paragraph_format.space_after = Pt(2)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    ending = '.' if last else ';'
    r = p.add_run(f'• {text}{ending}')
    r.font.size = Pt(11)
    if is_ph: r.font.color.rgb = RED


# Dieu 1 — Phe duyet phan tich thu cap
DIEU(1, [
    ('Phê duyệt việc phân tích thứ cấp và công bố quốc tế đối với '
     'bộ dữ liệu được thu thập trong khuôn khổ chương trình sàng '
     'lọc sức khỏe tâm thần học sinh thường niên của hai Trường '
     'Trung học cơ sở Nhật Tân và Tây Mỗ (Hà Nội), do nghiên cứu '
     'sinh Công Thị Hằng (Khoa Tâm lý - Giáo dục, Trường Đại học '
     'Sư phạm Hà Nội) thực hiện dưới sự hướng dẫn khoa học của ',
     False),
    ('[Học hàm/học vị]', True),
    (' Đào Minh Đức (', False),
    ('[Cơ quan công tác: Viện Tâm lý học Lâm sàng HOẶC Khoa Tâm lý '
     '- Giáo dục Trường ĐHSPHN — nhóm chọn]', True),
    ('). Đề cương đạt yêu cầu về tính khoa học và tuân thủ các '
     'nguyên tắc đạo đức trong nghiên cứu.', False),
])

# Dieu 2 — Pham vi phe duyet
DIEU(2, [('Phạm vi phê duyệt:', False)])

SUB('Tên đề tài luận án',
    [('"Các yếu tố nguy cơ và bảo vệ đối với các rối loạn lo âu ở '
      'học sinh trung học cơ sở Việt Nam"', False)])

SUB('Mục tiêu nghiên cứu',
    [('Xây dựng và kiểm chứng mô hình tích hợp ba yếu tố nguy cơ '
      '(bị bắt nạt, áp lực học tập, nghiện điện thoại) cùng bốn '
      'yếu tố bảo vệ (gắn bó trường, hỗ trợ từ cha mẹ, hỗ trợ từ '
      'bạn bè, lòng tự trọng) đến ba loại rối loạn lo âu theo '
      'DSM-5 (Lo âu Lan tỏa, Lo âu Xã hội, Lo âu Chia ly) ở học '
      'sinh trung học cơ sở Việt Nam', False)])

SUB('Đối tượng tham gia',
    [('Học sinh trung học cơ sở (tuổi 11-14, lớp 6 đến lớp 9)', False)])

SUB('Cỡ mẫu',
    [('1.352 học sinh', False)])

SUB('Địa bàn',
    [('Trường Trung học cơ sở Nhật Tân và Trường Trung học cơ sở '
      'Tây Mỗ, Hà Nội', False)])

SUB('Thời gian thu thập dữ liệu',
    [('năm học 2024-2025 (trong khuôn khổ chương trình sàng lọc '
      'sức khỏe tâm thần học sinh thường niên của nhà trường)', False)])

SUB('Phạm vi công bố kết quả',
    [('Luận án tiến sĩ của nghiên cứu sinh Công Thị Hằng và ba bài '
      'báo khoa học dự kiến công bố trên tạp chí quốc tế thuộc danh '
      'mục Scopus/WoS, tất cả khai thác cùng bộ dữ liệu nêu trên. '
      'Ba bài báo gồm: (i) Bài Q2 — "Mô hình phương trình cấu trúc '
      'tích hợp các yếu tố nguy cơ và bảo vệ đối với các phân loại '
      'rối loạn lo âu ở học sinh trung học cơ sở Việt Nam: Nghiên '
      'cứu cắt ngang"; (ii) Bài Q3 — "Các đường dẫn đặc thù theo '
      'giới đến các rối loạn lo âu ở giai đoạn đầu tuổi vị thành '
      'niên: Mô hình phương trình cấu trúc đa nhóm trên học sinh '
      'trung học cơ sở Việt Nam"; (iii) Bài Q4 — "Các hồ sơ phân '
      'loại lo âu ở học sinh trung học cơ sở Việt Nam: Phân tích '
      'hồ sơ tiềm ẩn tích hợp các chỉ báo nguy cơ và bảo vệ".',
      False)])

# Dieu 3 — Co so phap ly + chap thuan da co
DIEU(3, [('Cơ sở pháp lý và đạo đức của việc thu thập dữ liệu đã được '
          'thực hiện:', False)])
BULL('Việc thu thập dữ liệu được thực hiện trong khuôn khổ chương '
     'trình sàng lọc sức khỏe tâm thần học sinh thường niên của hai '
     'Trường Trung học cơ sở Nhật Tân và Tây Mỗ, theo thư đồng ý '
     'của Ban Giám hiệu hai trường (đính kèm hồ sơ)')
BULL('Sự đồng ý bằng văn bản (informed consent) của cha/mẹ/người '
     'giám hộ đã được thu thập trước khi học sinh tham gia khảo sát')
BULL('Sự đồng thuận bằng văn bản (assent) của học sinh đã được thu '
     'thập trước khi thực hiện khảo sát')
BULL('Quy trình thu dữ liệu tuân thủ Tuyên bố Helsinki (1964 và '
     'các sửa đổi tiếp theo) cùng các nguyên tắc đạo đức cơ bản: '
     'tự nguyện, bảo mật, không gây hại', last=True)

# Dieu 4 — Trach nhiem
DIEU(4, [('Trách nhiệm của nghiên cứu sinh và người hướng dẫn '
          'khoa học:', False)])
BULL('Đảm bảo bảo mật thông tin cá nhân của đối tượng nghiên cứu '
     'theo quy định hiện hành')
BULL('Báo cáo Hội đồng Đạo đức hàng năm về tiến độ phân tích và '
     'công bố')
BULL('Báo cáo kịp thời cho Hội đồng Đạo đức mọi tình huống bất lợi '
     'có thể phát sinh hoặc mọi thay đổi so với phạm vi đã được '
     'phê duyệt')
BULL('Lưu trữ hồ sơ nghiên cứu tối thiểu 05 năm kể từ ngày hoàn '
     'thành nghiên cứu')
BULL('Khi nộp bài báo cho tạp chí quốc tế, đính kèm bản scan '
     'Quyết định này (bản gốc tiếng Việt) và bản dịch tiếng Anh '
     'chính thức của Trường Đại học Sư phạm Hà Nội', last=True)

# Dieu 5 — Hieu luc
DIEU(5, [('Quyết định này có hiệu lực kể từ ngày ký. Phạm vi phê '
          'duyệt áp dụng trong toàn bộ thời gian triển khai phân '
          'tích và công bố kết quả nghiên cứu nêu tại Điều 2, đồng '
          'thời bao trùm mọi công bố khoa học phát sinh từ đề tài '
          'này.', False)])

# Dieu 6 — Noi nhan
DIEU(6, [('Nơi nhận: Nghiên cứu sinh Công Thị Hằng; Người hướng '
          'dẫn khoa học (', False),
         ('[Học hàm/học vị]', True),
         (' Đào Minh Đức); Khoa Tâm lý - Giáo dục, Trường Đại học '
          'Sư phạm Hà Nội; Phòng Quản lý Khoa học Công nghệ, Trường '
          'Đại học Sư phạm Hà Nội; Lưu Văn thư - Hội đồng Đạo đức.',
          False)])


# ============================================================
# Chu ky
# ============================================================
p = d.add_paragraph()
p.paragraph_format.space_before = Pt(14)
p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('CHỦ TỊCH HỘI ĐỒNG ĐẠO ĐỨC'); r.bold = True; r.font.size = Pt(11)

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('(Ký, ghi rõ họ tên và đóng dấu)')
r.italic = True; r.font.size = Pt(10)

for _ in range(3):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p.paragraph_format.space_after = Pt(2)

p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r = p.add_run('[Học hàm/học vị + Họ tên Chủ tịch HĐ Đạo đức]')
r.bold = True; r.font.size = Pt(11)
r.font.color.rgb = RED


# ============================================================
# Strip metadata
# ============================================================
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''; cp.keywords = ''
cp.comments = ''; cp.last_modified_by = ''; cp.category = ''
cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            chunks.append(cell.text)
total = sum(len(c.split()) for c in chunks)
print(f'WORD COUNT: {total}')
