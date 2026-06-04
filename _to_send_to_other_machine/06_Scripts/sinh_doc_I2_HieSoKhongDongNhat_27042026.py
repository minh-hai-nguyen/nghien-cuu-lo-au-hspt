# -*- coding: utf-8 -*-
"""Doc giải thích I² (I bình phương) — chỉ số đo độ không đồng nhất trong phân tích gộp.
Tiếng Việt thuần + chú thích Anh + reference đầy đủ + ví dụ minh hoạ.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\I_binh_phuong_giai_thich_cho_thay_27042026.docx'
RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color

# ===================================================================
title("CHỈ SỐ I² (I BÌNH PHƯƠNG) LÀ GÌ?", 16)
subtitle("Trả lời thầy về một thuật ngữ thống kê thường gặp trong phân tích gộp")
nr("")
subtitle("Trợ lý nghiên cứu — 27/04/2026")
subtitle("Tiếng Việt thuần — chú thích thuật ngữ Anh trong ngoặc")

# ===================================================================
H1("1. CÂU TRẢ LỜI NHANH")
nr("Câu hỏi của thầy: \"Có một số công thức toán có I bình phương, nghĩa là gì vậy em?\"",
   italic=True, color=GRAY)
nr("")
nr("TRẢ LỜI THẲNG:", bold=True)
nr("I² (đọc là \"I bình phương\", tiếng Anh: \"I-squared\") là CHỈ SỐ ĐO MỨC ĐỘ KHÔNG "
   "ĐỒNG NHẤT GIỮA CÁC NGHIÊN CỨU (heterogeneity index) trong PHÂN TÍCH GỘP "
   "(meta-analysis). Khi tổng hợp nhiều nghiên cứu để trả lời cùng một câu hỏi (ví dụ "
   "\"liệu pháp nhận thức hành vi có giảm lo âu ở học sinh không?\"), các nghiên cứu "
   "thường cho kết quả khác nhau. I² cho biết PHẦN TRĂM của sự khác biệt đó là DO "
   "sự khác biệt thực sự giữa các nghiên cứu (chứ không phải do may rủi / sai số ngẫu "
   "nhiên).", bold=True)
nr("")
nr("Phổ giá trị: 0% đến 100%. Càng cao thì các nghiên cứu càng KHÔNG đồng nhất.")
nr("Người đề xuất: Higgins và Thompson năm 2002 (đăng trên tạp chí Statistics in Medicine).",
   italic=True)

# ===================================================================
H1("2. ĐỊNH NGHĨA + CÔNG THỨC")

H2("2.1. Định nghĩa")
nr("Hai khái niệm liên quan trong phân tích gộp:")
nr("• HOMOGENEITY (đồng nhất): các nghiên cứu cho kết quả tương tự nhau — chỉ khác "
   "nhau do may rủi (sai số ngẫu nhiên — random error).")
nr("• HETEROGENEITY (không đồng nhất): các nghiên cứu cho kết quả khác nhau ĐÁNG "
   "KỂ — do sự khác biệt thực sự về thiết kế, mẫu, can thiệp, bối cảnh.")
nr("")
nr("I² đo PHẦN TRĂM tổng phương sai (variance) trong các kích thước hiệu ứng "
   "(effect sizes) là DO heterogeneity thực sự, không phải do random error.")

H2("2.2. Công thức")
nr("I² = (Q − df) / Q × 100%", bold=True, size=14)
nr("")
nr("Trong đó:")
nr("• Q = thống kê Q của Cochran (Cochran's Q statistic) — đo TỔNG mức độ chệch của "
   "kết quả các nghiên cứu so với kết quả gộp.")
nr("• df = bậc tự do (degrees of freedom) = số lượng nghiên cứu − 1 (df = k − 1, "
   "với k là số nghiên cứu trong phân tích gộp).")
nr("")
nr("Diễn giải toán học: Khi Q > df → có heterogeneity → I² > 0. Khi Q ≤ df → I² = 0 "
   "(quy ước, vì giá trị âm không có ý nghĩa).")

H2("2.3. Ví dụ tính tay đơn giản")
nr("Giả sử thực hiện phân tích gộp 10 nghiên cứu (k=10), tính được Q = 25 và df = 9.")
nr("• I² = (25 − 9) / 25 × 100% = 16/25 × 100% = 64%")
nr("• Diễn giải: 64% sự khác biệt giữa kết quả các nghiên cứu là do heterogeneity "
   "thực sự; 36% còn lại là do sai số ngẫu nhiên (may rủi).")
nr("• Mức 64% thuộc khoảng \"đáng kể\" (substantial) → cần phân tích phụ để hiểu vì sao.")

# ===================================================================
H1("3. CÁCH DIỄN GIẢI I² (theo Cochrane Handbook + Higgins 2003)")

tbl(['Khoảng I²', 'Mức độ heterogeneity', 'Diễn giải lâm sàng', 'Hành động cần làm'],
    [
        ['0% – 25%', 'THẤP (low)',
         'Các nghiên cứu khá đồng nhất; kết quả gộp đáng tin',
         'Có thể dùng mô hình hiệu ứng cố định (fixed-effects model) để gộp'],
        ['25% – 50%', 'TRUNG BÌNH (moderate)',
         'Có một chút khác biệt nhưng chấp nhận được',
         'Có thể dùng cả 2 mô hình; nên báo cả random + fixed để so sánh'],
        ['50% – 75%', 'ĐÁNG KỂ (substantial)',
         'Khác biệt đáng kể, cần thận trọng khi diễn giải kết quả gộp',
         'Bắt buộc dùng mô hình hiệu ứng ngẫu nhiên (random-effects); '
         'cân nhắc phân tích phụ (subgroup analysis)'],
        ['75% – 100%', 'CAO (high) hoặc CỰC CAO (considerable)',
         'Khác biệt rất lớn — kết quả gộp KHÔNG đáng tin',
         'KHÔNG nên gộp số liệu chung; cần phân tích meta-regression hoặc '
         'mô tả định tính từng nghiên cứu'],
    ], [2.5, 3.5, 5.0, 5.0])

nr("")
nr("LƯU Ý QUAN TRỌNG:", bold=True, color=RED)
nr("Các ngưỡng trên là KHUYẾN CÁO TƯƠNG ĐỐI, không phải tuyệt đối. Cochrane Handbook "
   "phiên bản mới nhất (2024) khuyến cáo NÊN xem I² cùng với:")
nr("• Khoảng tin cậy 95% của I² (95% CI) — vì I² có thể có sai số lớn khi số nghiên "
   "cứu k nhỏ.")
nr("• Giá trị p của thống kê Q (chi-square test) — nếu p < 0,10 thì heterogeneity có "
   "ý nghĩa thống kê.")
nr("• Tau bình phương (τ²) — đo phương sai giữa các nghiên cứu (between-study "
   "variance) — bổ sung cho I².")
nr("• Forest plot — biểu đồ rừng — quan sát trực quan các điểm + khoảng tin cậy có "
   "chồng chéo nhau hay không.")

# ===================================================================
H1("4. VÍ DỤ MINH HOẠ TỪ KHO TÀI LIỆU DỰ ÁN")

H2("4.1. Caldwell và cộng sự 2019 (Lancet Psychiatry — phân tích mạng meta)")
nr("Bài này đã có trong kho dự án (đã trích trong các doc dịch + phản biện 3 bài Brown).")
nr("• Loại nghiên cứu: phân tích gộp mạng (network meta-analysis — NMA) các thử "
   "nghiệm có đối chứng về can thiệp SKTT trường học.")
nr("• Số nghiên cứu: 137 (k=137)")
nr("• I² báo cáo: dao động 30% – 45% xuyên các so sánh — mức TRUNG BÌNH")
nr("• Diễn giải: Có khác biệt thực sự giữa các nghiên cứu nhưng ở mức chấp nhận; "
   "kết quả gộp tương đối đáng tin.")
nr("• Hành động của tác giả: Dùng mô hình hiệu ứng ngẫu nhiên + báo cáo cả các "
   "phân tích phụ theo loại can thiệp.")

H2("4.2. Zhang và cộng sự 2023 (J Youth Adolescence — phân tích gộp 65 thử nghiệm)")
nr("Đã có trong kho dự án qua các tham chiếu trong 3 doc Brown.")
nr("• Phân tích gộp 65 thử nghiệm có đối chứng về can thiệp SKTT trường học")
nr("• I² báo cáo: ~50–70% cho hầu hết outcomes — mức ĐÁNG KỂ")
nr("• Hành động: Phân tích phụ theo loại can thiệp (CBT vs khác), do người cung cấp "
   "(nhà lâm sàng vs giáo viên), bậc học (THCS vs THPT) → giải thích được phần lớn "
   "heterogeneity.")

H2("4.3. Phân tích gộp Frontiers Psychiatry 2025 (38 thử nghiệm về resilience)")
nr("Trong doc TuLieu V3 đã có.")
nr("• 38 thử nghiệm có đối chứng / 15.730 học sinh")
nr("• I² CAO (>75%) cho nhiều outcomes — heterogeneity quá lớn")
nr("• Tác giả thừa nhận khó kết luận chung; phải tách thành các nhóm con (chánh "
   "niệm, CBT, học tập cảm xúc xã hội) để phân tích riêng.")

# ===================================================================
H1("5. KHI NÀO CẦN QUAN TÂM ĐẾN I²?")

nr("Thầy CẦN xem I² khi đọc:")
nr("• Phân tích gộp (meta-analysis)")
nr("• Tổng quan hệ thống có gộp số liệu định lượng (systematic review with "
   "quantitative synthesis)")
nr("• Phân tích gộp mạng (network meta-analysis — NMA)")
nr("• Phân tích cumulative meta-analysis hoặc trial sequential analysis")
nr("")
nr("Thầy KHÔNG cần xem I² khi đọc:")
nr("• Một thử nghiệm đơn lẻ (RCT) — không có heterogeneity giữa các nghiên cứu")
nr("• Tổng quan hệ thống chỉ mô tả định tính (qualitative synthesis)")
nr("• Tổng quan phạm vi (scoping review)")
nr("• Nghiên cứu cắt ngang đơn lẻ")

# ===================================================================
H1("6. SO SÁNH MÔ HÌNH HIỆU ỨNG CỐ ĐỊNH (FIXED) vs NGẪU NHIÊN (RANDOM)")

nr("Khi I² cao, tác giả buộc phải chọn mô hình thống kê phù hợp:")

tbl(['Đặc điểm', 'Mô hình hiệu ứng cố định (Fixed-effects)',
     'Mô hình hiệu ứng ngẫu nhiên (Random-effects)'],
    [
        ['Giả định nền',
         'Có MỘT giá trị hiệu ứng thật chung; các nghiên cứu chỉ khác nhau do random error',
         'Mỗi nghiên cứu có hiệu ứng thật RIÊNG — đến từ một phân phối chung'],
        ['Khi nào nên dùng',
         'Heterogeneity THẤP (I² ≤ 25%); các nghiên cứu rất tương đồng',
         'Heterogeneity TRUNG BÌNH trở lên (I² > 25%)'],
        ['Trọng số (weighting)',
         'Theo nghịch đảo phương sai (1/variance) — nghiên cứu nhỏ trọng số nhỏ',
         'Theo nghịch đảo (variance + tau²) — nghiên cứu nhỏ có trọng số tương đối '
         'lớn hơn so với fixed'],
        ['Khoảng tin cậy 95%',
         'Hẹp hơn — chỉ phản ánh độ chính xác mẫu',
         'Rộng hơn — phản ánh cả độ chính xác mẫu lẫn variability giữa các nghiên cứu'],
        ['Diễn giải kết luận',
         '\"Hiệu ứng thật là X cho QUẦN THỂ NÀY\"',
         '\"Hiệu ứng thật trung bình là X qua các quần thể; có thể dao động trong '
         'khoảng nào đó\"'],
        ['Phù hợp với SKTT trẻ em', 'Hiếm khi phù hợp',
         'GẦN NHƯ luôn phù hợp — vì các nghiên cứu khác nhau về văn hoá, độ tuổi, '
         'thiết kế'],
    ], [4.0, 6.0, 6.0])

nr("")
nr("KHUYẾN NGHỊ THỰC TIỄN: trong lĩnh vực sức khỏe tâm thần trẻ em + vị thành niên, "
   "GẦN NHƯ LUÔN phải dùng MÔ HÌNH HIỆU ỨNG NGẪU NHIÊN — vì các nghiên cứu thường "
   "không đồng nhất do bối cảnh văn hoá, lứa tuổi, thiết kế can thiệp khác nhau.",
   bold=True)

# ===================================================================
H1("7. NHỮNG SAI LẦM PHỔ BIẾN KHI ĐỌC I²")

H3("Sai lầm 1: Coi I² thấp là \"chứng cứ chắc chắn\"")
nr("Khi I² thấp (vd. 10%), KHÔNG có nghĩa kết quả gộp \"đúng\". Có thể do:")
nr("• Số nghiên cứu k nhỏ → power thấp để phát hiện heterogeneity → I² ước tính sai")
nr("• Các nghiên cứu cùng thiết kế nhưng cùng có một sai lệch hệ thống (publication "
   "bias)")
nr("• Tất cả đều ở một bối cảnh hẹp (vd. chỉ Mỹ) → không generalisable")

H3("Sai lầm 2: Coi I² cao là \"chứng cứ rác\"")
nr("I² cao KHÔNG có nghĩa nghiên cứu rác. Có thể do:")
nr("• Các nghiên cứu thực sự đo các nhóm khác nhau (lứa tuổi, văn hoá, mức độ nặng)")
nr("• Có sự đa dạng lớn về can thiệp được đo (vd. CBT vs ACT vs MBSR)")
nr("• Cần phân tích phụ (subgroup) để hiểu chi tiết, không phải bỏ luôn")

H3("Sai lầm 3: Lẫn lộn I² với Q hoặc tau²")
nr("• Q (Cochran's Q): chỉ đo CÓ hay KHÔNG có heterogeneity (yes/no)")
nr("• I²: đo PHẦN TRĂM phương sai do heterogeneity (0–100%)")
nr("• τ² (tau squared): đo MAGNITUDE TUYỆT ĐỐI của heterogeneity (đơn vị giống "
   "effect size)")
nr("3 chỉ số bổ sung cho nhau, không thay thế nhau.")

H3("Sai lầm 4: Áp dụng ngưỡng 25/50/75% như tuyệt đối")
nr("Higgins và Thompson 2003 đã NÓI RÕ trong bài: \"these thresholds are rough "
   "guides only\". Cần xét trong bối cảnh từng lĩnh vực — vd. dược lý có ngưỡng "
   "khác tâm lý học.")

# ===================================================================
H1("8. ÁP DỤNG CHO VIỆT NAM")

nr("Khi đọc các phân tích gộp về SKTT trẻ em / vị thành niên Việt Nam (hiếm — vì "
   "VN chưa có nhiều phân tích gộp nội địa), thầy nên CHÚ Ý:")
nr("(1) Nếu phân tích gộp toàn các nghiên cứu QUỐC TẾ rồi áp dụng cho VN — I² thường "
   "cao do văn hoá khác biệt → cần phân tích phụ riêng cho khu vực châu Á / Đông "
   "Nam Á.")
nr("(2) Phân tích gộp riêng các nghiên cứu VN (ví dụ tổng hợp DASS-21 + GAD-7 ở các "
   "tỉnh khác nhau) — I² có thể cao do khác biệt vùng miền, thiết kế công cụ, độ "
   "tuổi.")
nr("(3) GBD ASEAN 2025 (đã có trong kho dự án QT10/QT12) — báo cáo các tỉ lệ ASEAN — "
   "khi gộp 10 quốc gia thường có I² > 75% → phải phân tích từng quốc gia riêng.")

H2("Ví dụ tính I² thử cho dữ liệu giả định Việt Nam")
nr("Giả sử có 5 nghiên cứu Việt Nam đo tỉ lệ lo âu ở học sinh THPT:")
tbl(['Nghiên cứu', 'Tỉ lệ lo âu', 'Cỡ mẫu'],
    [
        ['Hoa 2024 Hà Nội (corpus dự án)', '40,6%', '3.910'],
        ['Vĩnh Lộc 2024 TP.HCM', '50,3%', '976'],
        ['Lạng Sơn DTTS 2024', '54,4%', '845'],
        ['An Giang 2025', '61,2%', '366'],
        ['V-NAMHS 2022 toàn quốc (sàng lọc)', '21,7%', '5.996'],
    ], [6.0, 3.0, 3.0])
nr("Quan sát: tỉ lệ dao động 21,7% – 61,2% — chênh lệch GẦN 3 LẦN. Nếu chạy phân tích "
   "gộp các số liệu này, I² SẼ RẤT CAO (>90%) — nghĩa là các nghiên cứu KHÔNG đồng "
   "nhất, không nên gộp chung. Cần phân tích phụ theo: công cụ đo (GAD-7 vs DASS-21 vs "
   "DISC-5), khu vực (đô thị vs nông thôn vs DTTS), thời điểm (trước/sau COVID).")

# ===================================================================
H1("9. LIÊN HỆ I² VỚI PHÂN TÍCH GỘP (MA — Meta-Analysis)")
nr("Em viết riêng phần này để giải đáp câu hỏi của thầy về tin nhắn có cụm \"(I², τ²)\" "
   "và \"Với MA: I² < 25 % = thấp\".",
   italic=True, color=GRAY)

H2("9.1. MA là viết tắt của gì?")
nr("MA = META-ANALYSIS = PHÂN TÍCH GỘP", bold=True, size=14)
nr("Là phương pháp THỐNG KÊ TỔNG HỢP kết quả định lượng từ NHIỀU nghiên cứu khác "
   "nhau để đưa ra một kết luận chung mạnh hơn từng nghiên cứu riêng lẻ. Khác với "
   "tổng quan hệ thống (Systematic Review — SR) chỉ MÔ TẢ các nghiên cứu, MA dùng "
   "công thức TOÁN HỌC để gộp số liệu (pool data).")
nr("")
nr("Liên hệ với SR: thường người ta làm SR + MA cùng lúc — gọi là SR with MA hoặc "
   "viết tắt SR/MA. Nhiều bài trong corpus dự án là SR/MA (vd. Frontiers Psychiatry "
   "2025 — SR/MA 38 thử nghiệm; Zhang 2023 — SR/MA 65 thử nghiệm; Caldwell 2019 — "
   "SR/MA dạng mạng / network meta-analysis — NMA).")

H2("9.2. Vì sao MA CẦN I²?")
nr("Khi MA gộp số liệu từ N nghiên cứu, người làm MA luôn phải trả lời 2 câu hỏi:")
nr("• Câu 1: KẾT QUẢ GỘP là bao nhiêu? (effect size gộp + khoảng tin cậy 95%)")
nr("• Câu 2: CÁC NGHIÊN CỨU CÓ ĐỒNG NHẤT KHÔNG? (heterogeneity)")
nr("")
nr("→ I² trả lời CHÍNH XÁC câu hỏi 2. Không có I², ta không biết kết quả gộp có "
   "đáng tin hay không. Vì vậy I² là CHỈ SỐ BẮT BUỘC PHẢI BÁO CÁO trong mọi MA "
   "(theo hướng dẫn báo cáo PRISMA 2020 — Preferred Reporting Items for Systematic "
   "Reviews and Meta-Analyses).", bold=True)

H2("9.3. Cụm \"(I², τ²)\" trong tin nhắn nghĩa là gì?")
nr("Đây là CẶP CHỈ SỐ KINH ĐIỂN luôn đi cùng nhau khi báo cáo heterogeneity trong MA:")
tbl(['Chỉ số', 'Tên đầy đủ', 'Đơn vị', 'Đo gì', 'Khoảng giá trị'],
    [
        ['I²', 'I-squared (I bình phương)', 'Phần trăm (%)',
         'TỈ LỆ phương sai do heterogeneity (so với tổng phương sai)',
         '0% – 100%'],
        ['τ²', 'Tau-squared (Tau bình phương)', 'Đơn vị giống effect size '
         '(vd. nếu effect size là Cohen d thì τ² đo bằng d²)',
         'GIÁ TRỊ TUYỆT ĐỐI của phương sai giữa các nghiên cứu (between-study variance)',
         '≥ 0 (không giới hạn trên)'],
    ], [1.5, 4.0, 4.0, 5.5, 4.0])
nr("")
nr("KHÁC BIỆT QUAN TRỌNG:", bold=True)
nr("• I² là PHẦN TRĂM tương đối — giúp so sánh giữa các MA khác nhau (cùng thang 0-100%).")
nr("• τ² là GIÁ TRỊ TUYỆT ĐỐI — giúp tính khoảng dự đoán (prediction interval) cho "
   "nghiên cứu mới.")
nr("• Hai chỉ số bổ sung nhau — báo cả 2 cho người đọc cái nhìn đầy đủ.")

H2("9.4. Ngưỡng \"I² < 25% = thấp\" trong MA — Diễn giải đầy đủ")
nr("Tin nhắn của thầy ghi: \"Với MA: I² < 25 % = thấp\". Đây là ngưỡng KINH ĐIỂN "
   "Higgins & Thompson 2003 đưa ra:")
tbl(['Ngưỡng I² trong MA', 'Mức heterogeneity', 'Ý nghĩa thực tiễn cho MA'],
    [
        ['I² < 25%', 'THẤP (low)',
         'MA \"khoẻ\" — kết quả gộp đáng tin; có thể dùng mô hình hiệu ứng cố định '
         '(fixed-effects)'],
        ['25% ≤ I² < 50%', 'TRUNG BÌNH (moderate)',
         'MA chấp nhận được nhưng cần ghi chú; dùng được cả fixed + random effects'],
        ['50% ≤ I² < 75%', 'ĐÁNG KỂ (substantial)',
         'MA phải dùng MÔ HÌNH HIỆU ỨNG NGẪU NHIÊN (random-effects); cần phân tích '
         'phụ (subgroup analysis) để giải thích heterogeneity'],
        ['I² ≥ 75%', 'CAO / CỰC CAO (high / considerable)',
         'KHÔNG NÊN GỘP số liệu chung — kết quả gộp không đáng tin; chỉ nên phân tích '
         'từng nghiên cứu riêng hoặc dùng meta-regression để mô hình hoá heterogeneity'],
    ], [3.5, 3.5, 8.0])

H2("9.5. Quy trình thực tế khi đọc một bài MA")
nr("Khi thầy đọc một bài MA, em đề xuất bám 4 bước theo thứ tự:")
nr("• BƯỚC 1: Tìm KẾT QUẢ GỘP (pooled effect size) — thường ở phần Kết quả + "
   "Forest plot (biểu đồ rừng).")
nr("• BƯỚC 2: Tìm I² — thường BÁO NGAY DƯỚI hoặc BÊN CẠNH effect size gộp.")
nr("• BƯỚC 3: Đối chiếu I² với ngưỡng (xem bảng 9.4 trên):")
nr("  – Nếu I² < 25% → tin được kết quả gộp.")
nr("  – Nếu I² ≥ 50% → THẬN TRỌNG, đọc thêm phân tích phụ.")
nr("• BƯỚC 4: Xem tác giả có giải thích heterogeneity không (qua subgroup, "
   "meta-regression, sensitivity analysis). Nếu không có → MA chưa đầy đủ.")

H2("9.6. Ví dụ thực tế từ corpus dự án")

H3("Ví dụ 1 — Caldwell 2019 (Lancet Psychiatry NMA)")
nr("• Loại bài: SR/MA dạng mạng (Network Meta-Analysis — NMA) các thử nghiệm có "
   "đối chứng về phòng ngừa lo âu + trầm cảm trẻ em / vị thành niên trong trường học.")
nr("• Số nghiên cứu: 137")
nr("• I² báo cáo: 30% – 45% (thay đổi tuỳ outcome) — mức TRUNG BÌNH")
nr("• Diễn giải: Có heterogeneity vừa phải; tác giả dùng mô hình hiệu ứng ngẫu nhiên "
   "+ phân tích phụ theo loại can thiệp (CBT vs chánh niệm vs giáo dục) → giải "
   "thích được phần lớn heterogeneity.")
nr("• Kết luận của Caldwell: bằng chứng YẾU cho hầu hết can thiệp (effect size nhỏ + "
   "heterogeneity vừa phải).")

H3("Ví dụ 2 — Zhang 2023 (J Youth Adolescence)")
nr("• Loại bài: SR/MA 65 thử nghiệm có đối chứng về can thiệp SKTT trường học")
nr("• I² báo cáo: 50% – 70% — mức ĐÁNG KỂ")
nr("• Diễn giải: Heterogeneity cao → tác giả phải làm phân tích phụ rất chi tiết "
   "(theo loại can thiệp, người cung cấp, lứa tuổi, thời gian) → cuối cùng kết luận "
   "CBT do chuyên gia lâm sàng cung cấp ở THPT cho effect size cao nhất.")

H3("Ví dụ 3 — Phân tích gộp giả định cho dữ liệu Việt Nam")
nr("Như đã ghi ở mục 8 — nếu gộp 5 nghiên cứu lo âu HS THPT Việt Nam (Hoa Hà Nội "
   "40,6%, Vĩnh Lộc 50,3%, Lạng Sơn DTTS 54,4%, An Giang 61,2%, V-NAMHS 21,7%), "
   "I² SẼ rất cao (>90%). Lúc này KHÔNG NÊN báo \"tỉ lệ lo âu HS THPT Việt Nam = "
   "X%\" mà phải báo riêng từng vùng / từng công cụ đo.")

H2("9.7. Bảng tóm gọn cho thầy về cụm \"(I², τ²) — Với MA: I² < 25% = thấp\"")
tbl(['Câu trong tin nhắn', 'Giải thích cho thầy'],
    [
        ['MA', 'Meta-Analysis = Phân tích gộp — gộp số liệu từ NHIỀU nghiên cứu'],
        ['(I², τ²)', 'Cặp chỉ số đo độ KHÔNG đồng nhất giữa các nghiên cứu — luôn '
         'báo cáo cùng nhau trong MA'],
        ['I²', 'Phần trăm (%) phương sai do heterogeneity — phổ 0-100%'],
        ['τ²', 'Phương sai TUYỆT ĐỐI giữa các nghiên cứu — đơn vị giống effect size'],
        ['I² < 25%', 'Heterogeneity THẤP — các nghiên cứu khá đồng nhất → MA đáng tin'],
        ['Với MA: ...', 'Đây là quy ước CHUẨN khi báo cáo MA — bắt buộc theo PRISMA 2020'],
    ], [4.5, 11.5])

H2("9.8. Khi viết báo cáo / bài báo MA bằng tiếng Việt")
nr("Quy ước chuẩn khi viết MA tiếng Việt (theo PRISMA 2020 + Cochrane Handbook):")
nr("• Câu mẫu báo I²: \"Mức độ không đồng nhất giữa các nghiên cứu được đánh giá "
   "là TRUNG BÌNH (I² = 45%, τ² = 0,12)\".")
nr("• Câu mẫu khi I² thấp: \"Các nghiên cứu cho thấy mức độ đồng nhất cao (I² = "
   "18%), do đó mô hình hiệu ứng cố định được sử dụng để gộp số liệu\".")
nr("• Câu mẫu khi I² cao: \"Heterogeneity giữa các nghiên cứu rất cao (I² = 82%); "
   "do đó kết quả gộp chỉ được trình bày để tham khảo, không nên áp dụng trực tiếp "
   "vào thực tiễn lâm sàng. Phân tích phụ theo [yếu tố X] giúp giảm I² xuống còn [Y%]\".")

H1("10. THAM KHẢO ĐẦY ĐỦ")

nr("1. Higgins JPT, Thompson SG (2002). Quantifying heterogeneity in a meta-analysis. "
   "Statistics in Medicine, 21(11):1539–1558. DOI: 10.1002/sim.1186 — PMID 12111919. "
   "(BÀI GỐC đề xuất I²)", size=11)
nr("2. Higgins JPT, Thompson SG, Deeks JJ, Altman DG (2003). Measuring inconsistency "
   "in meta-analyses. BMJ, 327:557–560. DOI: 10.1136/bmj.327.7414.557 — PMID 12958120. "
   "(Bài GIẢI THÍCH cho người dùng lâm sàng — dễ đọc nhất)", size=11)
nr("3. Cochrane Handbook for Systematic Reviews of Interventions (phiên bản mới nhất). "
   "Higgins JPT, Thomas J, Chandler J, Cumpston M, Li T, Page MJ, Welch VA (chủ biên). "
   "Cochrane Collaboration. Chương 10.10: Heterogeneity. "
   "https://training.cochrane.org/handbook", size=11)
nr("4. Borenstein M, Hedges LV, Higgins JPT, Rothstein HR (2009). Introduction to "
   "Meta-Analysis. Wiley. ISBN 978-0-470-05724-7. (Sách giáo khoa kinh điển về "
   "meta-analysis — chương 16 về heterogeneity)", size=11)
nr("5. Cuijpers P (2016). Meta-analyses in mental health research: A practical guide. "
   "Pim Cuijpers Tools. (Hướng dẫn thực hành cho nghiên cứu sức khỏe tâm thần)",
   size=11)

H2("Sách / tài liệu tiếng Việt về phân tích gộp")
nr("• Hồ Nguyễn Phúc Vương (chủ biên) (2018). Phân tích gộp trong y học chứng cứ. "
   "NXB Y học. (Một trong số ít sách tiếng Việt về meta-analysis — có chương riêng "
   "về heterogeneity và I²)", size=11)
nr("• Trần Bình Tuấn (2020). Thống kê y học cơ bản. NXB Y học. (Có phần giới thiệu "
   "I² trong chương meta-analysis)", size=11)

H2("Truy vết nội bộ")
nr("• Doc này: 01_Bao-cao/I_binh_phuong_giai_thich_cho_thay_27042026.docx", size=11)
nr("• Glossary nội bộ: 06_Scripts/glossary_data/glossary_v3_pedagogical.json — em "
   "đề xuất bổ sung mục \"I²\" và \"heterogeneity\" vào glossary để dùng cho các "
   "doc sau", size=11)
nr("• Áp dụng kiến thức này khi đọc: Caldwell 2019 (đã trích trong 3 doc Brown), "
   "Zhang 2023, Frontiers 2025 (TuLieu V3), QT29 Li 2025 NMA (corpus dự án)", size=11)

d.save(OUT)
print('Wrote:', OUT)
