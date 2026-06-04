# -*- coding: utf-8 -*-
"""Sinh 2 file ho tro NCS chinh sua luan an trong 75 phut:
  (A) Phan hoi tiep thu chinh sua - bang Response to Reviewer
  (B) Doan giai trinh thay doi mau khao sat - file rieng
26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUTDIR = os.path.join(ROOT, 'Luận án TS')

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)

def P(doc, text, indent=True, italic=False, color=None, bold=False):
    """Tao paragraph - tu dong DETECT [...] de TO DO + BOLD (khong highlight)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    # Split text bang pattern [...] de to do cac placeholder
    parts = re.split(r'(\[[^\]]+\])', text)
    for part in parts:
        if not part:
            continue
        r = p.add_run(part)
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        if part.startswith('[') and part.endswith(']'):
            # Day la placeholder [...] - to do dam
            r.font.color.rgb = RGBColor(192, 0, 0)
            r.bold = True
        else:
            r.italic = italic; r.bold = bold
            if color is not None:
                r.font.color.rgb = color
    return p

def P_cell(cell_para, text):
    """Tao text trong cell - tu dong to do [...]."""
    parts = re.split(r'(\[[^\]]+\])', text)
    for part in parts:
        if not part:
            continue
        r = cell_para.add_run(part)
        r.font.name = 'Times New Roman'; r.font.size = Pt(11)
        if part.startswith('[') and part.endswith(']'):
            r.font.color.rgb = RGBColor(192, 0, 0)
            r.bold = True

# ============================================================
# FILE A: PHAN HOI TIEP THU CHINH SUA
# ============================================================
doc_a = doc_init()
H(doc_a, 'BẢN GIẢI TRÌNH TIẾP THU CHỈNH SỬA LUẬN ÁN', 1)
P(doc_a, '(Phản hồi nhận xét của PGS Đặng Hoàng Minh, ngày 25/05/2026)', indent=False, italic=True)
doc_a.add_paragraph()
P(doc_a, 'NCS xin được giải trình các nội dung đã chỉnh sửa theo nhận xét lần 2 của thầy phản biện như sau. Các điểm chính được tóm tắt ở bảng dưới đây; các điểm chi tiết cùng vị trí cụ thể trong luận án được nêu ở phụ lục đi kèm.', indent=True)

# Bang chinh
headers = ['Mục phản biện', 'Nội dung phản biện', 'NCS đã chỉnh sửa']
rows = [
    ['Mục 2 — Thay đổi khách thể',
     'Quan ngại về việc thay đổi toàn bộ khách thể (1.352 HS THCS Nhật Tân + Tây Mỗ) trong 2 tháng, ảnh hưởng đến đề cương, hội đồng đạo đức, tiến độ.',
     'Đã bổ sung đoạn giải trình quy trình điều chỉnh khách thể ở Mở đầu (trang [điền số trang]), trong đó nêu rõ lý do điều chỉnh, văn bản chấp thuận của cơ sở đào tạo và Hội đồng đạo đức số [điền số], ngày [điền ngày].'],
    ['Mục 3 — Cách trình bày tiếp thu',
     'NCS cần nêu rõ "đã chỉnh sửa thế nào" thay vì chỉ "đã viết lại".',
     'Đã bổ sung bảng giải trình tiếp thu (chính là bản này) ở đầu luận án, kèm tham chiếu trang/mục cho từng điểm chỉnh sửa.'],
    ['Mục 10 — Cách viết tổng quan',
     'NCS liệt kê từng công trình thay vì tổng quát hóa theo chủ đề; nhận xét về công cụ (1.1.1.3, 1.1.2.3) không khớp công cụ thực tế được điểm luận.',
     'Đã viết lại toàn bộ mục 1.1 theo cấu trúc nhóm chủ đề (Tỷ lệ – Mức độ – Biểu hiện / Yếu tố nguy cơ / Yếu tố bảo vệ / Can thiệp), với các nghiên cứu Việt Nam gom vào cuối mỗi chủ đề (xem trang [điền]). Phần nhận xét công cụ đã được rà soát lại để chỉ đề cập đến các công cụ thực tế có trong điểm luận.'],
    ['Mục 8 + 14 + 28 + 49 — Biến "biện pháp đối phó"',
     'Biến "biện pháp đối phó" được trình bày không nhất quán: lúc là yếu tố bảo vệ, lúc là biến phụ thuộc.',
     'Đã làm rõ "biện pháp đối phó" được nghiên cứu như một biến trung gian (mediator) trong mối quan hệ giữa yếu tố nguy cơ và biểu hiện rối loạn lo âu. Đã chỉnh sửa cách trình bày ở cơ sở lý luận (trang [điền]), phương pháp (trang [điền]) và kết luận (trang [điền]) để bảo đảm nhất quán.'],
    ['Mục 15 đến 19 — Phân loại RLLA, DSM',
     'Phân loại RLLA dưới các góc độ chưa đúng cấu trúc; dùng lẫn lộn DSM-IV, DSM-5, DSM-5-TR; "RLLA tổng quát" vs "lan tỏa".',
     'Đã thống nhất sử dụng DSM-5 trong toàn luận án; đã sửa "DSM V" thành "DSM-5"; đã thống nhất thuật ngữ "rối loạn lo âu lan tỏa" theo bản tiếng Việt chính thức của DSM-5. Đã rà soát lại cấu trúc phân loại ở mục 1.2.2.'],
    ['Các mục còn lại (sâu)',
     'Các critique về phân tích SEM (mục 38-39), factor loading (mục 32-34), bổ sung phương pháp quan sát + phỏng vấn sâu (mục 9, 25, 30), restructure mục 3.7 (mục 41-46), Bảng 3.21-3.22 (mục 37).',
     'Các điểm này đòi hỏi điều chỉnh sâu hơn về thiết kế nghiên cứu và phân tích thống kê. NCS xin được làm việc thêm với thầy hướng dẫn và cập nhật trong lần nộp tiếp theo.'],
]

t = doc_a.add_table(rows=1 + len(rows), cols=3)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
widths = [3.5, 6.0, 6.5]
set_grid(t, widths)
for row in t.rows:
    for ci in range(3):
        colw(row.cells[ci], widths[ci])

# Header
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    shade(c, 'D9E2F3')

# Rows - dung P_cell de auto-highlight [...]
for ri, rd in enumerate(rows):
    for ci, v in enumerate(rd):
        c = t.rows[ri + 1].cells[ci]
        # Xoa cell mac dinh paragraph cu (chua text "")
        c.text = ''
        cell_para = c.paragraphs[0]
        cell_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        P_cell(cell_para, v)

doc_a.add_paragraph()
P(doc_a, 'NCS xin chân thành cảm ơn thầy phản biện đã dành thời gian đọc và đưa ra các nhận xét quý báu. Các góp ý của thầy giúp luận án được hoàn thiện hơn về cả cấu trúc, nội dung khoa học và phương pháp luận. Những điểm chưa kịp xử lý trong lần này sẽ được NCS tiếp tục làm việc cùng thầy hướng dẫn để bổ sung trong các lần nộp tới.', indent=True)
doc_a.add_paragraph()
P(doc_a, 'Hà Nội, ngày [điền] tháng [điền] năm 2026', indent=False, italic=False)
P(doc_a, 'Nghiên cứu sinh', indent=False, bold=True)
P(doc_a, '', indent=False)
P(doc_a, '', indent=False)
P(doc_a, 'Công Thị Hằng', indent=False, bold=True)

out_a = os.path.join(OUTDIR, '02_PhanHoi_TiepThu_ChinhSua_v1_26052026.docx')
doc_a.save(out_a)
print(f"File A: {out_a}")
print(f"  Size: {os.path.getsize(out_a) // 1024}KB")

# ============================================================
# FILE B: DOAN GIAI TRINH THAY DOI MAU
# ============================================================
doc_b = doc_init()
H(doc_b, 'ĐOẠN GIẢI TRÌNH QUY TRÌNH ĐIỀU CHỈNH KHÁCH THỂ NGHIÊN CỨU', 1)
P(doc_b, '(Soạn để chèn vào phần Mở đầu, sau mục Lý do chọn đề tài — phục vụ phản hồi Mục 2 phản biện của PGS Đặng Hoàng Minh)', indent=False, italic=True)
doc_b.add_paragraph()

H(doc_b, 'Bản 1 — Phiên bản đầy đủ (~180 từ)', 2)
P(doc_b, 'HƯỚNG DẪN: Trước khi sử dụng, NCS điền các thông tin trong [ngoặc vuông] bằng thông tin thực tế. Bỏ phần [ngoặc vuông] khi đã điền xong.', indent=False, italic=True, color=RGBColor(192, 0, 0))
P(doc_b, '', indent=False)
P(doc_b,
  'Trong quá trình triển khai luận án, sau khi nộp bản thảo lần 1 vào tháng 3/2026 và tiếp thu các nhận xét của thầy phản biện, nghiên cứu sinh đã trao đổi với thầy hướng dẫn và xác định cần điều chỉnh khách thể nghiên cứu nhằm bảo đảm tính đại diện và phù hợp với mục tiêu của đề tài. Cụ thể, mẫu nghiên cứu mới gồm 1.352 học sinh trung học cơ sở thuộc hai trường Trung học cơ sở Nhật Tân và Trung học cơ sở Tây Mỗ tại Hà Nội — đại diện cho khu vực [nội thành – ngoại thành / chất lượng cao – đại trà / điền đặc điểm vùng]. Việc điều chỉnh khách thể đã được [Hội đồng đạo đức nghiên cứu của cơ sở đào tạo / Phòng Đào tạo sau đại học] thông qua tại [văn bản số …, ngày …]. Toàn bộ quy trình thu thập dữ liệu mới đều tuân thủ các nguyên tắc đạo đức nghiên cứu trên trẻ em – vị thành niên: có sự đồng ý của Ban Giám hiệu nhà trường, sự đồng thuận của phụ huynh và bản thân học sinh tham gia khảo sát, bảo mật thông tin cá nhân và quyền rút khỏi nghiên cứu bất cứ lúc nào. Đề cương chi tiết và tiến độ luận án sau điều chỉnh cũng đã được báo cáo lại với cơ sở đào tạo và được chấp thuận tại [văn bản số …, ngày …].',
  indent=True)
doc_b.add_paragraph()

H(doc_b, 'Bản 2 — Phiên bản ngắn gọn (~110 từ)', 2)
P(doc_b, 'HƯỚNG DẪN: Dùng bản này nếu không muốn chiếm quá nhiều dòng ở Mở đầu. Cũng cần điền [ngoặc vuông].', indent=False, italic=True, color=RGBColor(192, 0, 0))
P(doc_b, '', indent=False)
P(doc_b,
  'Sau khi tiếp thu nhận xét phản biện lần 1 (tháng 3/2026) và trao đổi với thầy hướng dẫn, nghiên cứu sinh đã điều chỉnh khách thể nghiên cứu để bảo đảm phù hợp với mục tiêu đề tài. Mẫu mới gồm 1.352 học sinh THCS thuộc trường THCS Nhật Tân và THCS Tây Mỗ, Hà Nội. Việc điều chỉnh đã được [cơ sở đào tạo / Hội đồng đạo đức] thông qua tại [văn bản số …, ngày …]; quy trình thu thập dữ liệu mới tuân thủ các nguyên tắc đạo đức nghiên cứu trên vị thành niên (đồng ý của nhà trường, đồng thuận của phụ huynh và học sinh, bảo mật thông tin, quyền rút khỏi nghiên cứu).',
  indent=True)
doc_b.add_paragraph()

H(doc_b, 'Bản 3 — Phiên bản tối giản (~60 từ)', 2)
P(doc_b, 'HƯỚNG DẪN: Dùng nếu chỉ muốn nói ngắn gọn trong 1-2 câu. Đặt ở footnote hoặc cuối đoạn Lý do chọn đề tài.', indent=False, italic=True, color=RGBColor(192, 0, 0))
P(doc_b, '', indent=False)
P(doc_b,
  'Khách thể nghiên cứu trong luận án này (1.352 HS tại trường THCS Nhật Tân và THCS Tây Mỗ, Hà Nội) đã được điều chỉnh từ mẫu của bản thảo lần 1, sau khi tiếp thu phản biện và được [cơ sở đào tạo / Hội đồng đạo đức] thông qua tại [văn bản số …, ngày …].',
  indent=True)
doc_b.add_paragraph()

H(doc_b, 'Lưu ý quan trọng', 2)
P(doc_b, 'Đoạn giải trình này CHỈ có giá trị nếu việc điều chỉnh khách thể thực sự đã được cơ sở đào tạo / Hội đồng đạo đức thông qua bằng văn bản. Nếu chưa có văn bản chính thức, NCS cần làm rõ với thầy hướng dẫn trước khi sử dụng đoạn này — không nên ghi thông tin không có thật vào luận án.', indent=True, italic=True, color=RGBColor(192, 0, 0))
P(doc_b, 'Nếu chưa có văn bản chính thức nhưng đã có sự chấp thuận miệng từ thầy hướng dẫn và đại diện cơ sở đào tạo, NCS có thể viết theo hướng:', indent=True, italic=True)
P(doc_b, '"Việc điều chỉnh khách thể đã được nghiên cứu sinh thực hiện theo định hướng của thầy hướng dẫn và đang hoàn tất các thủ tục báo cáo chính thức với cơ sở đào tạo."', indent=True, italic=True, color=RGBColor(192, 0, 0))

out_b = os.path.join(OUTDIR, '03_GiaiTrinh_ThayDoi_KhachThe_v1_26052026.docx')
doc_b.save(out_b)
print(f"\nFile B: {out_b}")
print(f"  Size: {os.path.getsize(out_b) // 1024}KB")

print(f"\nDa tao 2 file. Tong thoi gian uoc tinh de chinh sua tay khi co 2 file nay: 75 phut.")
