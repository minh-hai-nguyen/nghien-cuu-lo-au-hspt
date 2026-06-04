# -*- coding: utf-8 -*-
"""Fix LA v3_2 -> v3_3:
- 'M. Sakia' -> 'A.M. Saikia' (spelling sai, para 174)
- 'Xu và cs (2022)' / 'Xu và cs. (2022)' -> 'Xu và cs. (2021)' (year mismatch with TLTK 1545)
- 'Xu và cs (2021)' giu nguyen
27/05/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_2_FixSoLieu_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_3_FixSaikiaXu_27052026.docx')
# Force re-run by always re-doing from v3_2

RED = RGBColor(192, 0, 0)

REPLACEMENTS = [
    # Saikia spelling
    ('M. Sakia và cs.(2023)', 'A.M. Saikia và cs. (2023)'),
    ('M. Sakia và cs. (2023)', 'A.M. Saikia và cs. (2023)'),
    ('M. Sakia và cs.', 'A.M. Saikia và cs.'),
    # Xu year - 2022 -> 2021 (TLTK 1545 cites 2021)
    ('Q. Xu và cs (2022)', 'Q. Xu và cs. (2021)'),
    ('Q. Xu và cs. (2022)', 'Q. Xu và cs. (2021)'),
    ('Xu và cs (2022)', 'Xu và cs. (2021)'),
    ('Xu và cs. (2022)', 'Xu và cs. (2021)'),
    ('(Q. Xu và cs., 2022)', '(Q. Xu và cs., 2021)'),
    # Qingqing Xu intro paragraph year fix
    ('Chongjian Wang, Cuiping Wu (2022)', 'Chongjian Wang, Cuiping Wu (2021)'),
]


def apply_replacements_in_paragraph(p, replacements, applied_tracker, idx):
    full = p.text
    if not full:
        return False
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            i = full.find(old, cursor)
            if i == -1: break
            matches.append((i, i + len(old), old, new))
            cursor = i + len(old)
    if not matches:
        return False
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    safe = []
    last = -1
    for m in matches:
        if m[0] >= last:
            safe.append(m)
            last = m[1]
    segments = []
    cur = 0
    for s, e, o, n in safe:
        if cur < s:
            segments.append((full[cur:s], False))
        segments.append((n, True))
        applied_tracker.append(f"para {idx}: '{o}' -> '{n}'")
        cur = e
    if cur < len(full):
        segments.append((full[cur:], False))
    for r in p.runs:
        r.text = ''
    for text, is_red in segments:
        if not text: continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


print(f"=== Fix LA v3_2 -> v3_3 ===")
backup = SRC.replace('.docx', '_BACKUP_v32_27052026.docx')
if not os.path.exists(backup):
    shutil.copy(SRC, backup)

d = Document(SRC)
applied = []
for i, p in enumerate(d.paragraphs):
    apply_replacements_in_paragraph(p, REPLACEMENTS, applied, i)
for ti, tb in enumerate(d.tables):
    for ri, row in enumerate(tb.rows):
        for ci, cell in enumerate(row.cells):
            for cpi, p in enumerate(cell.paragraphs):
                apply_replacements_in_paragraph(p, REPLACEMENTS, applied,
                                                 f'T{ti+1}r{ri}c{ci}p{cpi}')
d.save(OUT)
print(f"Saved: {OUT}")
print(f"Total: {len(applied)} fixes")
for a in applied:
    print(f"  ✓ {a}")
