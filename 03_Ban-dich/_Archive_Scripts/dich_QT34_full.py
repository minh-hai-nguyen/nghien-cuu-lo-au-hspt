# -*- coding: utf-8 -*-
"""Dịch đầy đủ QT34 — Korea MH Trends 2006-2022 — Cho et al. — Nature Sci Rep"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link bài báo gốc: https://doi.org/10.1038/s41598-024-XXXXX', size=10)

add_heading(doc, 'Xu hướng quốc gia về sức khỏe tâm thần thanh thiếu niên theo mức thu nhập tại Hàn Quốc, trước và sau COVID-19, 2006–2022', 1)
h = doc.add_paragraph()
r = h.add_run("National Trends in Adolescents' Mental Health by Income Level in South Korea, Pre- and Post-COVID-19, 2006-2022")
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', "National Trends in Adolescents' Mental Health by Income Level in South Korea, Pre- and Post-COVID-19, 2006-2022"),
    ('Tiêu đề dịch', 'Xu hướng quốc gia về sức khỏe tâm thần thanh thiếu niên theo mức thu nhập tại Hàn Quốc, trước và sau COVID-19, 2006–2022'),
    ('Tác giả', 'Jaehyeong Cho, Jaeyu Park, Hayeon Lee, Hyesu Jo, Sooji Lee, Hyeon Jin Kim, Yejun Son, Hyunjee Kim, Selin Woo, Seokjun Kim, Jiseung Kang, Damiano Pizzol, Jiyoung Hwang, Lee Smith, Dong Keon Yon'),
    ('Cơ quan', 'Center for Digital Health, Medical Science Research Institute, Kyung Hee University; Department of Pediatrics, Kyung Hee University Medical Center, Seoul; Nhiều cơ quan khác tại Hàn Quốc, Ý, Anh'),
    ('Tạp chí', 'Nature Scientific Reports (Q1, IF ≈ 4,6)'),
    ('Thông tin xuất bản', '2024, 14 trang, 4 bảng, 1 hình'),
    ('DOI', '10.1038/s41598-024-XXXXX'),
    ('Loại NC', 'Phân tích xu hướng 17 năm (2006–2022) — dữ liệu khảo sát quốc gia'),
    ('Mẫu', '1.138.804 VTN Hàn Quốc (tuổi TB = 15,01; SD = 0,75; 51,57% nam) từ KYRBS'),
    ('Nguồn dữ liệu', 'KYRBS — Korea Youth Risk Behavior Web-based Survey (Bộ Giáo dục + KCDC)'),
])
add_page_ref(doc, '1–14', 'Nature Scientific Reports', '2024')

# TÓM TẮT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Mặc dù tác động đáng kể của đại dịch COVID-19 lên nhiều yếu tố liên quan SKTT VTN như stress, buồn bã, ý tưởng tự tử và cố gắng tự tử, NC về chủ đề này vẫn chưa đủ.')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Dựa trên KYRBS (Korea Youth Risk Behavior Web-based Survey) từ 2006 đến 2022. Phân tích vấn đề SKTT VTN dựa trên bảng hỏi + phỏng vấn y tế, trong 5 nhóm thu nhập, so sánh với nhiều yếu tố nguy cơ. Tổng 1.138.804 người tham gia (tuổi TB = 15,01; SD = 0,75; 51,57% nam). Phân tích hồi quy xu hướng, chia 8 giai đoạn (5 trước COVID + 3 trong/sau COVID).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
r = p.add_run('Kết quả: Năm 2022, tỷ lệ stress ở nhóm thu nhập CAO nhất: 40,07% (KTC 95%: 38,67–41,48), buồn bã 28,15%, ý tưởng tự tử 13,92%, cố tự tử 3,42%. Ở nhóm thu nhập THẤP nhất: stress 62,77% (59,42–66,13), buồn bã 46,83%, ý tưởng tự tử 31,70%, cố tự tử 10,45%. Nhóm thu nhập thấp có tỷ lệ cao hơn đáng kể với nhiều yếu tố nguy cơ. Tổng thể: xu hướng GIẢM trước COVID → TĂNG đáng kể trong/sau COVID. Bất bình đẳng thu nhập MỞ RỘNG trong đại dịch.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_p(doc, 'Kết luận: NC cho thấy mối liên quan giữa mức thu nhập hộ gia đình và tỷ lệ bệnh tật SKTT ở VTN. COVID-19 đã làm trầm trọng hơn bệnh tật SKTT ở VTN từ hộ thu nhập thấp, nhấn mạnh nhu cầu tăng cường giám sát và hỗ trợ SKTT.')

# ĐÁNH GIÁ NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Nature Scientific Reports Q1 IF ≈ 4,6 — tạp chí uy tín.',
    'KYRBS quốc gia — MẪU CỰC LỚN: 1.138.804 VTN, 17 năm (2006–2022). Dài nhất trong tất cả NC của Đề tài.',
    'Phân tầng theo THU NHẬP (5 nhóm) — ÍT NC nào làm. Chênh 22,7 điểm stress giữa giàu/nghèo.',
    'Mô hình ĐẢO CHIỀU: SKTT CẢI THIỆN trước COVID (2006–2019) → XẤU ĐI sau COVID — DUY NHẤT trong Đề tài (phương Tây tăng liên tục).',
    'Bất bình đẳng thu nhập MỞ RỘNG trong COVID — nhóm nghèo bị ảnh hưởng nặng hơn.',
    'Nhiều yếu tố nguy cơ phân tích: giới, rượu, thuốc lá, trình độ cha mẹ, thành tích học tập, vùng.',
    'Tự tử: ý tưởng 13,92% nhóm giàu vs 31,70% nhóm nghèo — chênh GẤP 2,3 lần.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Chỉ Hàn Quốc — hệ thống giáo dục cạnh tranh cực cao (suneung), văn hóa khác VN.',
    'Đo stress/buồn bã/tự tử bằng câu hỏi đơn — KHÔNG tách lo âu riêng (không GAD-7, DASS-21).',
    'KYRBS tự báo cáo trực tuyến — thiên lệch tự báo cáo.',
    'Thu nhập gia đình tự báo cáo bởi VTN — có thể không chính xác.',
    'Sinh thái — xu hướng cấp quốc gia, không cá nhân.',
]:
    add_p(doc, f'• {b}')

# 1. GIỚI THIỆU
add_page_ref(doc, '1–3', 'Nature Scientific Reports', '2024')
add_heading(doc, '1. GIỚI THIỆU', 2)
add_p(doc, 'Mặc dù tác động lớn của COVID-19 lên nhiều yếu tố liên quan SKTT VTN, NC về chủ đề này vẫn chưa đủ. Hầu hết NC tập trung vào stress học tập hoặc vấn đề SKTT riêng của VTN, nhưng cần điều tra sâu hơn về yếu tố kinh tế xã hội — đặc biệt bất bình đẳng thu nhập hộ gia đình.')
add_p(doc, 'Từ khi COVID-19 bùng phát và phong tỏa, cá nhân trải qua stress và lo âu tăng do sợ lây nhiễm và gián đoạn sinh hoạt. Nhiều người đối mặt cô lập xã hội, cô đơn, giảm tiếp cận dịch vụ trị liệu — dẫn đến giảm cả sức khỏe thể chất và tâm thần. Tác động tiêu cực đặc biệt rõ ở VTN, với tăng đáng kể các chỉ số SKTT tiêu cực.')
add_p(doc, 'VTN là giai đoạn quan trọng cho phát triển cá nhân, hình thành nhân cách. Trạng thái SKTT trong giai đoạn này không chỉ định hình giá trị cốt lõi mà còn có tác động lâu dài lên quỹ đạo cuộc sống. Do đó, đảm bảo sức khỏe SKTT VTN trong những năm then chốt là thiết yếu cho chuyển đổi lành mạnh sang tuổi trưởng thành.')
add_p(doc, 'Hàn Quốc có hệ thống giáo dục cạnh tranh cực cao — kỳ thi suneung (대학수학능력시험) quyết định tương lai, tạo áp lực khổng lồ. Tỷ lệ tự tử VTN cao nhất OECD. KYRBS (Korea Youth Risk Behavior Web-based Survey) được triển khai hàng năm từ 2005 bởi KCDC và Bộ Giáo dục, lấy mẫu cụm nhiều bậc đại diện quốc gia cho HS THCS + THPT.')
add_p(doc, 'Hàn Quốc phản ứng COVID khác nhiều nước: không phong tỏa toàn diện (lockdown), thay vào đó dùng truy vết + cách ly + xét nghiệm quy mô lớn. Tuy nhiên, vẫn có đóng cửa trường và chuyển sang học trực tuyến. Tỷ lệ tử vong chỉ 0,13% — thấp nhất trong 30 nước có ca nhiễm cao nhất. Giai đoạn COVID chia 3: 2020, 2021, 2022.')

# 2. PHƯƠNG PHÁP
add_page_ref(doc, '2–3', 'Nature Scientific Reports', '2024')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)
add_p(doc, '2.1. Nguồn dữ liệu', bold=True)
add_p(doc, 'KYRBS: Khảo sát Hành vi Nguy cơ Thanh niên Hàn Quốc, tự báo cáo trực tuyến, hàng năm từ 2005. Lấy mẫu cụm nhiều bậc — đại diện quốc gia cho HS THCS (lớp 7–9) + THPT (lớp 10–12). Tổng 1.138.804 người tham gia 2006–2022. Tuổi TB = 15,01 (SD = 0,75). 587.256 nam (51,57%).')

add_p(doc, '2.2. Biến số', bold=True)
add_p(doc, '4 chỉ số SKTT (câu hỏi đơn):')
add_p(doc, '• Mức stress nhận thức: "Bạn thường trải qua bao nhiêu stress?" (1–5, phân loại: thấp/trung bình/cao).')
add_p(doc, '• Buồn bã: "Trong 12 tháng qua, bạn có trải qua buồn bã hoặc tuyệt vọng nghiêm trọng đến mức gián đoạn hoạt động hàng ngày liên tục 2 tuần?" (Có/Không).')
add_p(doc, '• Ý tưởng tự tử: "Trong 12 tháng qua, bạn có nghiêm túc nghĩ đến tự tử?" (Có/Không).')
add_p(doc, '• Cố gắng tự tử: "Trong 12 tháng qua, bạn có cố gắng tự tử?" (Có/Không).')
add_p(doc, 'Thu nhập hộ gia đình: 5 nhóm (cao, trung cao, trung, trung thấp, thấp) — tự báo cáo bởi VTN.')
add_p(doc, 'Yếu tố phân tầng: giới, lớp, vùng cư trú, rượu gần đây, hút thuốc, trình độ cha mẹ, thành tích học tập.')

add_p(doc, '2.3. Phân tích', bold=True)
add_p(doc, 'Hồi quy xu hướng — β (95% CI). 8 giai đoạn: Trước COVID (2006–2008, 2009–2011, 2012–2014, 2015–2017, 2018–2019) + Trong/Sau COVID (2020, 2021, 2022). So sánh xu hướng trước vs sau COVID: βdiff (khác biệt β). Trọng số dân số — ước tính tỷ lệ có trọng số.')

# 3. KẾT QUẢ
add_page_ref(doc, '3–8', 'Nature Scientific Reports', '2024')
add_heading(doc, '3. KẾT QUẢ', 2)
add_p(doc, '3.1. Tổng quan mẫu', bold=True)
add_p(doc, 'Tổng 1.138.804 VTN, 17 năm. Phân bố đều theo giới (51,57% nam). Đại diện quốc gia cho HS THCS+THPT Hàn Quốc.')

add_heading(doc, 'Bảng 1. Xu hướng stress nhận thức cao theo thu nhập 2006–2022', 3)
add_table(doc,
    ['Thu nhập', '2006–08', '2012–14', '2018–19', '2020', '2022', 'β trước COVID', 'β sau COVID', 'βdiff'],
    [['Cao nhất', '38,93%', '33,93%', '34,81%', '28,30%', '40,07%', '−1,41***', '2,39***', '3,80***'],
     ['Trung cao', '40,92%', '35,11%', '37,27%', '31,75%', '38,82%', '−1,21***', '0,93***', '2,14***'],
     ['Trung', '43,50%', '38,70%', '39,65%', '33,56%', '40,41%', '−1,46***', '0,59*', '2,05***'],
     ['Trung thấp', '53,90%', '49,56%', '52,06%', '45,82%', '52,10%', '−1,11***', '0,30 (n.s.)', '1,41***'],
     ['Thấp nhất', '63,37%', '60,22%', '59,44%', '52,85%', '62,77%', '−1,40***', '0,98 (n.s.)', '2,38**'],
     ['Chênh Giàu-Nghèo', '24,44đ', '26,29đ', '24,63đ', '24,55đ', '22,70đ', '', '', '']],
    widths=[2.0, 1.5, 1.5, 1.5, 1.5, 1.5, 2.0, 2.0, 1.5])
add_p(doc, '*** p < 0,001; ** p < 0,01; * p < 0,05. β = thay đổi trung bình hàng năm (%/năm). βdiff = khác biệt xu hướng trước vs sau COVID.', size=9, italic=True)

add_heading(doc, 'Bảng 2. Tỷ lệ 4 chỉ số SKTT năm 2022 theo thu nhập', 3)
add_table(doc,
    ['Thu nhập', 'Stress cao (%)', 'Buồn bã (%)', 'Ý tưởng tự tử (%)', 'Cố tự tử (%)', 'Chênh vs Cao nhất'],
    [['Cao nhất', '40,07', '28,15', '13,92', '3,42', 'Reference'],
     ['Trung cao', '38,82', '28,68', '15,05', '3,34', ''],
     ['Trung', '40,41', '30,14', '15,94', '3,54', ''],
     ['Trung thấp', '52,10', '40,64', '21,21', '5,49', '+12đ stress, +7đ tự tử'],
     ['Thấp nhất', '62,77', '46,83', '31,70', '10,45', '+22,7đ stress, GẤP 2,3 tự tử']],
    widths=[2.5, 2.0, 2.0, 2.5, 2.0, 4.0])
add_p(doc, 'Bất bình đẳng RẤT LỚN: Stress nhóm nghèo 62,77% vs giàu 40,07% (chênh 22,7 điểm). Cố tự tử: 10,45% vs 3,42% (GẤP 3 lần).', size=9, italic=True)

add_p(doc, '3.2. Mô hình đảo chiều', bold=True)
add_p(doc, 'Phát hiện then chốt: SKTT VTN Hàn Quốc CẢI THIỆN trước COVID (β < 0 cho tất cả chỉ số) → XẤU ĐI sau COVID (β > 0 hoặc thay đổi đáng kể). Đây là mô hình ĐẢO CHIỀU — KHÁC BIỆT với phương Tây (tăng liên tục):')
add_p(doc, '• Na Uy (QT21): tăng liên tục 13 năm (2011–2024).')
add_p(doc, '• Mỹ (QT23): lo âu TĂNG GẤP ĐÔI (2013–2021, AOR = 2,17).')
add_p(doc, '• Ireland (QT32): tăng trước COVID (2012–2019).')
add_p(doc, '• GBD (QT30): AAPC 0,84% toàn cầu (1990–2021).')
add_p(doc, '• Hàn Quốc (bài này): GIẢM trước COVID → TĂNG sau — DUY NHẤT đảo chiều.')
add_p(doc, 'Gợi ý: (1) Can thiệp SKTT CÓ THỂ hiệu quả (giai đoạn cải thiện 2006–2019), (2) COVID đảo ngược tiến bộ, đặc biệt ở nhóm nghèo.')

add_heading(doc, 'Bảng 3. So sánh xu hướng Hàn Quốc với các nước trong Đề tài', 3)
add_table(doc,
    ['Nước', 'Giai đoạn', 'Xu hướng', 'Đặc thù', 'Nguồn'],
    [['Hàn Quốc', '2006–2022 (17 năm)', 'Giảm → TĂNG (đảo chiều)', 'Bất bình đẳng thu nhập mở rộng', 'Bài này (QT34)'],
     ['Hoa Kỳ', '2013–2021 (8 năm)', 'Lo âu TĂNG GẤP ĐÔI', 'AOR = 2,17, chẩn đoán lâm sàng', 'JAACAP 2024 (QT23)'],
     ['Na Uy', '2011–2024 (13 năm)', 'Tăng liên tục', 'Trường + MXH giải thích', 'Norway 2025 (QT21)'],
     ['Ireland', '2012–2019 (7 năm)', 'Tăng (trước COVID)', 'Nữ nhanh hơn; OGA bảo vệ', 'Ireland 2024 (QT32)'],
     ['Úc', '2018–2022 (4 năm)', 'Flourishing giảm', 'CMH toàn diện', 'EpiPsychSci 2025 (QT25)'],
     ['Anh', '2017–2023', 'Lo âu gấp đôi (7–16)', '12%→20%; NHS £16 tỷ', 'UK NHS 2024 (QT26)'],
     ['Toàn cầu', '1990–2021 (31 năm)', 'AAPC 0,84%', '204 nước, GBD', 'GBD 2025 (QT30)'],
     ['VN', '2021 vs 2023', '41,5% → 25,4%', 'Phục hồi sau COVID?', 'Hoàng Trung Học (VN14)']],
    widths=[2.0, 3.0, 3.0, 3.0, 2.5])

add_p(doc, '3.3. Yếu tố nguy cơ', bold=True)
add_p(doc, 'Yếu tố nguy cơ mạnh nhất cho SKTT tiêu cực: (1) NỮ giới, (2) Rượu, (3) Hút thuốc. Nhóm thu nhập thấp có tỷ lệ cao hơn với nhiều yếu tố nguy cơ — bất lợi tích lũy.')

add_heading(doc, 'Bảng 4. Bất bình đẳng thu nhập và SKTT', 3)
add_table(doc,
    ['Thu nhập', 'Stress 2022', 'Buồn bã 2022', 'Tự tử 2022', 'So với VN'],
    [['Cao nhất', '40,07%', '28,15%', '13,92%', ''],
     ['Thấp nhất', '62,77%', '46,83%', '31,70%', ''],
     ['Chênh lệch', '22,70 điểm', '18,68 điểm', '17,78 điểm', 'VN: CHƯA CÓ dữ liệu theo thu nhập'],
     ['Xu hướng chênh', 'MỞ RỘNG sau COVID', '', '', 'Gap #4 cross-study']],
    widths=[2.5, 2.5, 2.5, 2.5, 4.0])

# 4. THẢO LUẬN
add_page_ref(doc, '4–6', 'Nature Scientific Reports', '2024')
add_heading(doc, '4. THẢO LUẬN', 2)
add_p(doc, 'Đây là NC đầu tiên theo dõi dài hạn 17 năm mối quan hệ thu nhập hộ gia đình — SKTT VTN, cùng xu hướng. Phát hiện chính:')
add_p(doc, '(1) Thu nhập thấp liên quan mạnh với stress, buồn bã, ý tưởng tự tử, cố tự tử. Có thể do cảm giác thiếu thốn tương đối so với bạn bè — xu hướng xã hội đặt nặng giàu có. VTN nghèo cảm thấy tách biệt, dẫn đến tự trọng thấp, suy nghĩ và hành vi trầm cảm.')
add_p(doc, '(2) COVID-19 làm trầm trọng bất bình đẳng: phong tỏa → HS nghèo sống chật chội, ít riêng tư, khó tiếp cận học trực tuyến → stress và cô lập tăng. Gia đình nghèo mất thu nhập → bất ổn tài chính → SKTT xấu đi nhanh hơn nhóm giàu.')
add_p(doc, '(3) Mô hình đảo chiều: SKTT cải thiện 2006–2019 — có thể do chính sách SKTT học đường Hàn Quốc (Wee Center trong trường, tư vấn VTN). COVID đảo ngược tất cả tiến bộ.')
add_p(doc, '(4) Nữ giới: tỷ lệ cao hơn nam cho tất cả chỉ số — phù hợp xu hướng toàn cầu: 59 Countries (QT31) AOR = 1,51, Ireland (QT32) nữ tăng nhanh hơn, Úc (QT25) flourishing nữ 36,3% vs nam 52,2%.')
add_p(doc, '(5) Tỷ lệ tự tử VTN Hàn Quốc cao nhất OECD — dữ liệu xác nhận xu hướng đáng lo ngại: ý tưởng tự tử tăng sau COVID (đặc biệt nhóm nghèo 31,70%).')

# 5. HẠN CHẾ
add_heading(doc, '5. ĐIỂM MẠNH VÀ HẠN CHẾ', 2)
add_p(doc, 'Điểm mạnh:', bold=True)
add_p(doc, '• KYRBS quốc gia — lấy mẫu cụm, đại diện toàn quốc.')
add_p(doc, '• N = 1.138.804 — mẫu lớn nhất trong Đề tài (cùng JAACAP 13,7 triệu).')
add_p(doc, '• 17 năm liên tục — xu hướng dài nhất.')
add_p(doc, '• Phân tầng thu nhập 5 nhóm — ít NC nào làm.')
add_p(doc, '• So sánh trước/sau COVID bằng βdiff.')

add_p(doc, 'Hạn chế:', bold=True)
add_p(doc, '• Chỉ Hàn Quốc — hệ thống giáo dục cạnh tranh cực cao (suneung), văn hóa khác VN. Tuy nhiên, áp lực học tập cũng rất cao ở VN → có thể tương đồng.')
add_p(doc, '• Đo stress/buồn bã/tự tử bằng CÂU HỎI ĐƠN — KHÔNG dùng thang chuẩn hóa (GAD-7, DASS-21, PHQ-9). Không tách lo âu riêng. Ireland (QT32) dùng DASS-21; Norway (QT21) dùng HSCL-6.')
add_p(doc, '• Tự báo cáo trực tuyến — VTN có thể không báo cáo chính xác thu nhập gia đình và tự tử (nhạy cảm, kỳ thị).')
add_p(doc, '• Sinh thái — xu hướng cấp quốc gia. Không theo dõi cùng cá nhân (cắt ngang lặp lại, không cohort).')
add_p(doc, '• Phản ứng COVID Hàn Quốc (không lockdown toàn diện) khác nhiều nước → tác động COVID có thể khác ở VN.')

# 6. KẾT LUẬN
add_heading(doc, '6. KẾT LUẬN', 2)
add_p(doc, 'Dữ liệu KYRBS 1.138.804 VTN Hàn Quốc qua 17 năm (2006–2022) cho thấy:')
add_p(doc, '(1) Thu nhập thấp liên quan mạnh với tất cả chỉ số SKTT tiêu cực — chênh lệch stress 22,7 điểm (62,77% nghèo vs 40,07% giàu), cố tự tử GẤP 3 lần.')
add_p(doc, '(2) Mô hình ĐẢO CHIỀU: SKTT cải thiện 2006–2019 → xấu đi 2020–2022 — COVID đảo ngược tiến bộ.')
add_p(doc, '(3) Bất bình đẳng MỞ RỘNG trong COVID — nhóm nghèo chịu tác động nặng hơn.')
add_p(doc, '(4) Nữ, rượu, thuốc lá — yếu tố nguy cơ mạnh nhất.')
add_p(doc, 'Hàm ý cho VN: Cần phân tích SKTT VTN VN theo thu nhập — V-NAMHS 2022 chưa phân tầng kinh tế đầy đủ. Nếu VN có mô hình tương tự Hàn Quốc, nhóm VTN nghèo/DTTS (Ngô Anh Vinh 2024: DTTS 54,4%) có thể đang bị ảnh hưởng nặng nhất.')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
refs = [
    'Cho, J., Park, J., Lee, H., et al. (2024). National trends in adolescents\' mental health by income level in South Korea, pre- and post-COVID-19, 2006-2022. Scientific Reports.',
    'Korea Centers for Disease Control and Prevention. Korea Youth Risk Behavior Web-based Survey (KYRBS), 2006–2022.',
    'Keyes, K.M., Gary, D., O\'Malley, P.M., et al. (2019). Recent increases in depressive symptoms among US adolescents. J Adolescent Health, 65, 590–598.',
    'Von Soest, T., Wichstrøm, L. (2014). Secular trends in depressive symptoms among Norwegian adolescents. J Abnorm Child Psychol, 42, 403–415.',
    'Cosma, A., et al. (2020). Cross-national time trends in adolescent mental well-being from 2002 to 2018. JAMA, 66(S), S50–S58.',
]
for ref in refs:
    add_p(doc, ref, size=10)
add_p(doc, '(Xem đầy đủ trong bài gốc — 34 TLTK)', size=10, italic=True)

# VIẾT TẮT
add_abbreviation_table(doc, [
    ('KYRBS', 'Korea Youth Risk Behavior Web-based Survey — Khảo sát Hành vi Nguy cơ Thanh niên Hàn Quốc'),
    ('KCDC', 'Korea Centers for Disease Control and Prevention'),
    ('β', 'Beta — Hệ số xu hướng (thay đổi %/năm)'),
    ('βdiff', 'Beta difference — Chênh lệch xu hướng trước/sau COVID'),
    ('KTC', 'Khoảng Tin cậy (Confidence Interval)'),
    ('OECD', 'Organisation for Economic Co-operation and Development'),
    ('SKTT', 'Sức khỏe tâm thần'),
    ('VTN', 'Vị thành niên'),
    ('DTTS', 'Dân tộc thiểu số'),
    ('MXH', 'Mạng xã hội'),
    ('OGA', 'One Good Adult — Một Người Lớn Tốt (Ireland 2024)'),
    ('CMH', 'Complete Mental Health — Sức khỏe Tâm thần Toàn diện (Keyes)'),
])

# PHẢN BIỆN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'Nature Scientific Reports Q1. KYRBS quốc gia 17 năm — xu hướng DÀI NHẤT trong Đề tài.',
    'N = 1.138.804 — mẫu lớn nhất (cùng JAACAP 13,7 triệu). Norway 2025 (QT21): 979K; 59 Countries (QT31): 179K.',
    'Phân tầng THU NHẬP — đóng góp ĐỘC ĐÁO. Ít NC nào phân tích bất bình đẳng kinh tế → SKTT VTN. Ireland (QT32): không phân tầng kinh tế; Úc (QT25): cũng không.',
    'Mô hình ĐẢO CHIỀU — phát hiện quan trọng: SKTT VTN CÓ THỂ cải thiện (giai đoạn 2006–2019 ở Hàn Quốc). Phương Tây tăng liên tục — có thể do chính sách SKTT Hàn Quốc hiệu quả (Wee Center).',
    'So sánh trước/sau COVID bằng βdiff — phương pháp thống kê rõ ràng.',
    'Tự tử phân tầng thu nhập — hiếm. Ý tưởng tự tử 31,70% nhóm nghèo vs 13,92% giàu — chênh GẤP 2,3 lần. 59 Countries (QT31): ý tưởng tự tử AOR = 2,84 cho lo âu.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'Chỉ Hàn Quốc — hệ thống giáo dục cạnh tranh (suneung), tỷ lệ tự tử VTN OECD cao nhất. Văn hóa confucian khác VN nhưng có tương đồng (áp lực học tập cao). Cần NC tương tự ở VN.',
    'KHÔNG tách LO ÂU riêng — đo stress/buồn bã/tự tử bằng câu hỏi đơn. Norway (QT21) dùng HSCL-6; Ireland (QT32) dùng DASS-21; JAACAP (QT23) dùng chẩn đoán. Thiếu so sánh trực tiếp với các NC dùng thang chuẩn.',
    'Thu nhập gia đình TỰ BÁO CÁO bởi VTN — có thể không chính xác. VTN có thể không biết thu nhập gia đình hoặc cảm thấy nhạy cảm. Cần dữ liệu thu nhập từ cha mẹ hoặc hồ sơ hành chính.',
    'KYRBS tự báo cáo TRỰC TUYẾN — thiên lệch tự báo cáo, đặc biệt cho câu hỏi nhạy cảm (tự tử). VTN có thể không báo cáo trung thực. V-NAMHS 2022: chỉ 5,1% phụ huynh nhận ra con cần giúp đỡ.',
    'Cắt ngang lặp lại — không theo dõi cùng cá nhân. Không xác lập nhân quả: thu nhập thấp → SKTT kém hay SKTT kém → học kém → thu nhập thấp?',
    'Phản ứng COVID Hàn Quốc ĐẶC BIỆT (không lockdown toàn diện, truy vết + xét nghiệm) — tác động COVID có thể KHÁC ở VN (phong tỏa nghiêm ngặt hơn ở một số giai đoạn).',
    'Không đo screen time/MXH — yếu tố quan trọng theo Norway 2025 (QT21: MXH giải thích xu hướng) và Nature 2025 (QT27).',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống nghiên cứu / Research Gap:', bold=True)
for s in [
    'VN CHƯA CÓ phân tích SKTT VTN theo thu nhập — GAP #4 cross-study. V-NAMHS 2022 thu thập SES nhưng chưa phân tích chi tiết.',
    'Cần GSHS VN phân tầng kinh tế — so sánh nhóm nghèo/giàu. Ngô Anh Vinh 2024 (VN15): DTTS 54,4% lo âu — có thể phản ánh bất bình đẳng kinh tế (DTTS = thu nhập thấp).',
    'Mô hình đảo chiều ở VN? Hoàng Trung Học 2025 (VN14): lo âu 41,5% (2021) → 25,4% (2023) — có thể giảm tương tự Hàn Quốc trước COVID. Cần theo dõi 2024–2025 xem có đảo chiều không.',
    'Chính sách SKTT học đường Hàn Quốc (Wee Center) — VN có thể học hỏi. Zhameden 2025 (QT03): 0 RCT can thiệp trường tại VN.',
    'So sánh bất bình đẳng SKTT: Hàn Quốc (5 nhóm thu nhập) vs VN (vùng miền: đô thị vs nông thôn vs DTTS). Hoa 2024 (VN01): chỉ Hà Nội đô thị; An Giang 2025 (VN18): nông thôn 61,2%.',
    'COVID và bất bình đẳng ở VN — phong tỏa TPHCM 2021 (nghiêm ngặt hơn Hàn Quốc) có thể tạo bất bình đẳng SKTT lớn hơn. Cần NC đánh giá.',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '34_Korea_MH_Trends_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
