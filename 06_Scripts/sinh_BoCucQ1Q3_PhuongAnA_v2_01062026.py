# -*- coding: utf-8 -*-
"""Sinh bo cuc chi tiet IMRaD cho 2 bai Q1 + Q3 - v2 sau confirm tu thay.

V2 fixes:
- 3 authors: Hang Thi Cong + Nguyen Minh Duc + Duc Minh Dao
- Q1 = BMC Psychiatry (confirmed - highest probability)
- Q3 = PLOS ONE (confirmed - highest probability)
- Ethics requirement section detailed
- All citation PDFs verified in 02_Papers-goc/
- Saikia 2023 REAL PDF found at 11_Saikia_2023_IJCM.pdf
- All numbers (F, p, M, beta, R²) cross-checked with LA + THUC TRANG source

01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'BoCuc_Q1_Q3_PhuongAnA_v2_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'
s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.3


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('• ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


def V(text):
    """Verified mark - green"""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('✓ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)


# ============================================================
# COVER
# ============================================================
H1('BỐ CỤC CHI TIẾT 2 BÀI BÁO Q1 + Q3 (v2)')
P('(Phương án A — Split theo phương pháp + câu hỏi nghiên cứu)', italic=True)
P('Cập nhật sau confirm từ thầy: 3 authors + BMC Psychiatry (Q1) + PLOS ONE (Q3)', italic=True)
P('Ngày soạn: 01/06/2026', italic=True)
P('')
P('Cơ sở dữ liệu: 1.352 HS THCS Hà Nội (614 nam / 738 nữ)', italic=True)
P('Nguồn: file "Thực trạng RLLA HS THCS" thầy gửi 01/06/2026 + LA chính '
  '(`01_LuanAn_v4_ChuanTrinhBay_28052026.docx`)', italic=True)

d.add_page_break()


# ============================================================
# CHANGE SUMMARY (v2 vs v1)
# ============================================================
H1('TÓM TẮT THAY ĐỔI v2 vs v1')
B('(1) Authorship — thêm Nguyễn Minh Đức làm đồng tác giả; thứ tự cuối: '
  'Hang Thi Cong (1st), Nguyen Minh Duc, Duc Minh Dao (corresponding)')
B('(2) Q1 chốt = BMC Psychiatry (IF 4.4, ~30% acceptance, accept mixed-methods, '
  'reputation sạch)')
B('(3) Q3 chốt = PLOS ONE (IF 3.7, ~50% acceptance, broad scope, accept '
  'descriptive)')
B('(4) Ethics section detailed — em đã giải thích IRB requirement; '
  'NCS cần xin retroactive IRB letter HNUE trước khi submit')
B('(5) **Saikia 2023 PDF tìm thấy** — `11_Saikia_2023_IJCM.pdf` là REAL Saikia, '
  'KHÔNG phải mất như em báo cáo hôm qua')
B('(6) Tất cả 17 cited papers đều có PDF verified trong `02_Papers-goc/`')
B('(7) Mọi β, R², F, p, M, SD trong outline đã verify khớp LA chính + nguồn '
  '"Thực trạng" thầy gửi')


d.add_page_break()


# ============================================================
# PART 1: BÀI Q1
# ============================================================
H1('PHẦN 1 — BÀI BÁO Q1')

H2('1.1 Thông tin chung')
B('Title (tentative): "Integrated risk-protective structural equation model of '
  'anxiety disorder subtypes among Vietnamese lower secondary school students: '
  'A mixed-methods study"')
B('**Target journal**: BMC Psychiatry (Q1, IF 4.4, acceptance ~30%) ← chốt')
B('Backup: BMC Psychology (IF 2.6) nếu BMC Psychiatry reject')
B('Word count: 6.000–8.000 (excluding references)')
B('**Authors byline**: Hang Thi Cong¹* | Nguyen Minh Duc² | Duc Minh Dao¹†')
B('(*: 1st author, †: corresponding author)', 1)

H2('1.2 Novel contributions (3 điểm mới)')
B('(1) **First integrated SEM** kết hợp đồng thời 3 yếu tố nguy cơ (academic '
  'pressure, smartphone addiction, school bullying) + 4 yếu tố bảo vệ (school '
  'engagement, parental support, peer support, self-esteem) → 3 dạng RLLA theo '
  'DSM-5 trên mẫu HS THCS Việt Nam')
B('(2) **Gender-invariance phát hiện** của separation anxiety (F=0,246; p=0,620) '
  'trái với GAD (F=44,484; p<0,001) và Social Anxiety (F=45,984; p<0,001) — '
  'thách thức giả định gender-difference đồng đều cho mọi anxiety subtype')
B('(3) **Mixed-methods integration**: dữ liệu định tính phỏng vấn (HS NND lớp 9, '
  'MNA lớp 9) làm phong phú interpretation của các β coefficients')

H2('1.3 Cấu trúc IMRaD')

H3('Title + Abstract (~250 từ — structured Abstract chuẩn BMC Psychiatry)')
B('Background: anxiety prevalence + Vietnam gap + integrated model rationale')
B('Methods: 1,352 lower secondary students; 8 validated scales; CFA → SEM → '
  'multi-group invariance')
B('Results: SEM fit indices; β coefficients risk + protective; gender invariance '
  'finding; key qualitative themes')
B('Conclusions: prevention implications; cultural context')

H3('Introduction (~1.200 từ — 4 paragraphs)')
B('§1 Burden anxiety in adolescents — cite Xu 2021 (N=373.216 Chinese), '
  'Anderson 2025 (narrative review), GBD ASEAN 2025')
B('§2 Risk-protective framework + Vietnam gap — cite Lazarus & Folkman 1984, '
  'Compas 2017 (meta N=80,850), V-NAMHS 2022 (2,3% prevalence)')
B('§3 Differential gender patterns hypothesis — McLean 2011, Wen 2020, '
  'Saikia 2023 (men > women paradox in Northeast India)')
B('§4 Present study + 3 hypotheses')

H3('Methods (~1.500 từ) — 8 SCALES (TẤT CẢ VERIFIED)')
B('2.1 Participants: **1.352 HS** từ 2 trường THCS Hà Nội (Nhật Tân + Tây Mỗ); '
  '**614 nam / 738 nữ**; khối 6-9')
B('2.2 Measures (8 scales — VERIFIED trong LA P957-962):', 0)
B('Anxiety: RCADS (Chorpita, 2000) — 3-factor: GAD 7 items, SAD 4, SocAD 4', 1)
B('Risk: ESSA (Sun et al., 2011, 4 items), SAS-SV (Kwon et al., 2013, 5 items), '
  'OBVQ (Olweus, 1996, 8 items)', 1)
B('Protective: PSSM (Goodenow, 1993, 7 items), MSPSS (Zimet et al., 1988, 8 items, '
  '2 subscales), RSES (Rosenberg, 1965, 5 items)', 1)
B('Coping: Brief COPE (Carver, 1997, 15 items, 3 factors)', 1)
B('2.3 Qualitative interviews: semi-structured (NCS sẽ confirm số lượng sau)')
B('2.4 Analytic strategy:', 0)
B('Cronbach α + McDonald ω + CFA per scale', 1)
B('Integrated SEM (AMOS 31.0): 7 latent predictors → 3 anxiety latent outcomes', 1)
B('Multi-group invariance by gender (configural → metric → scalar)', 1)
B('Mixed-methods joint display (Creswell & Plano Clark, 2018)', 1)
B('2.5 Ethics: HNUE IRB approval + parental consent + student assent + school '
  'permission (2 trường) — **NCS CẦN XIN RETROACTIVE IRB LETTER**')

H3('Results (~1.800 từ) — TẤT CẢ SỐ LIỆU VERIFIED')
B('3.1 Sample characteristics: 1.352 HS; 614 nam (45,4%); 738 nữ (54,6%); '
  'khối 6: 368, khối 7: 316, khối 8: 340, khối 9: 328')
B('3.2 Psychometric: CFA per scale với fit indices')
B('3.3 Descriptive anxiety levels:', 0)
B('GAD (RLLATQ): M=55,82 (SD=22,39) — RCADS 7-item rút gọn', 1)
B('Separation (RLLACL): M=25,06 (SD=24,29)', 1)
B('Social (RLLAXH): M=48,41 (SD=26,19)', 1)
B('3.4 SEM main model (verified vs LA Bảng 32-42):', 0)
B('Risk → Total anxiety: ALHT β=+0,533*** > NĐT β=+0,400*** > '
  'BNHĐ β=+0,276*** (R²=0,284)', 1)
B('Protective → Total anxiety: TTr β=-0,457*** > HTCM β=-0,273*** > '
  'GBTH β=-0,108*** > HTBB β=-0,015 ns (R²=0,209)', 1)
B('Subtype-specific β (highlight): Bullying → Separation Anxiety β=0,376*** '
  '(strongest of any single path to SAD)', 1)
B('3.5 Multi-group invariance by gender — KEY FINDING:', 0)
B('GAD: female M=59,47 > male M=51,43; F=44,484; p<0,001', 1)
B('Social: female M=52,74 > male M=43,20; F=45,984; p<0,001', 1)
B('**Separation: gender-invariant** — F=0,246; p=0,620 — full invariance', 1)
B('3.6 Mixed-methods: qualitative themes confirm/extend β')

H3('Discussion (~1.700 từ)')
B('4.1 Summary 3 hypotheses + gender invariance')
B('4.2 Risk factors: academic pressure cultural specificity (Pascoe 2020 SR; '
  'Steare 2023 SR — 48/52 studies confirm)')
B('4.3 Self-esteem ~85-89% magnitude of academic pressure (verified LA P1416)')
B('4.4 Separation anxiety gender-invariance — NOVEL: cultural collectivism '
  'Vietnamese; developmental task universal across genders pre-puberty')
B('4.5 Clinical implications: Khung CT 8 nội dung từ LA')
B('4.6 Limitations: cross-sectional, 2 schools Hà Nội only, self-report')
B('4.7 Future directions: longitudinal, multi-site VN, RCT intervention')

H3('References (~40-50 refs) — VERIFIED PDFs')
B('Empirical: Xu 2021, Chen 2023, Wen 2020, Anderson 2025, Pascoe 2020, '
  'Steare 2023, Jefferies 2020, Nakie 2022, V-NAMHS 2022, Hoang Trung Hoc 2025, '
  'Bhardwaj 2020, Qiu 2022, Zhu 2025, Alharbi 2019, **Saikia 2023** '
  '(verified PDF: 11_Saikia_2023_IJCM.pdf)')
B('Methodological: Compas 2017 (meta-analysis), Hu & Bentler 1999 (SEM fit), '
  'Braun & Clarke 2006 (thematic), Creswell & Plano Clark 2018 (mixed-methods)')
B('Theoretical: Lazarus & Folkman 1984, Carver 1997, Rosenberg 1965, '
  'Goodenow 1993, Zimet 1988, Chorpita 2000, Sun 2011, Kwon 2013, Olweus 1996, '
  'McLean 2011, Bronfenbrenner & Morris 2006')

H3('Tables + Figures plan')
B('Table 1: Sample demographics (1.352 HS)')
B('Table 2: Psychometric properties (α, ω, CFA fit) — 8 scales')
B('Table 3: SEM β coefficients (risk + protective → 3 anxiety subtypes)')
B('Table 4: Multi-group invariance fit indices')
B('Table 5: Qualitative themes joint display')
B('Figure 1: Integrated SEM diagram')
B('Figure 2: Multi-group gender comparison')

d.add_page_break()


# ============================================================
# PART 2: BÀI Q3
# ============================================================
H1('PHẦN 2 — BÀI BÁO Q3')

H2('2.1 Thông tin chung')
B('Title (tentative): "Manifestations and patterns of anxiety disorder subtypes '
  'among Vietnamese lower secondary school students: A cross-sectional '
  'descriptive study"')
B('**Target journal**: PLOS ONE (Q1/Q2 by SJR, IF 3.7, acceptance ~50%) ← chốt')
B('Backup: BMC Pediatrics (IF 2.5) nếu PLOS muốn pediatric-specific')
B('Word count: 3.500–5.000')
B('**Authors byline**: Hang Thi Cong¹* | Nguyen Minh Duc² | Duc Minh Dao¹†')

H2('2.2 Novel contributions (3 điểm — phù hợp Q3/wide-scope tier)')
B('(1) **First detailed item-level analysis** RCADS Vietnamese adaptation (rút gọn '
  '21→15 items) trên mẫu HS THCS lớn N=1.352')
B('(2) **Cross-sectional grade trajectory** (khối 6→9) cho 3 dạng RLLA — '
  '**separation anxiety GIẢM** liên tục (32,13 → 27,14 → 20,88 → 19,46) phù hợp '
  'DSM-5 developmental literature')
B('(3) **Item-level identification** "lo lắng khi nghĩ mình đã không làm tốt..." '
  'M=64,28 — cung cấp screening targets cụ thể cho VN')

H2('2.3 Cấu trúc IMRaD (đơn giản hơn Q1)')

H3('Title + Abstract (~200 từ)')
B('Background: anxiety in adolescents + Vietnam under-researched item-level')
B('Methods: 1,352 students; RCADS Vietnamese; descriptive + ANOVA')
B('Results: M/SD per item; demographic + grade comparisons')
B('Conclusions: screening implications')

H3('Introduction (~800 từ — 3 paragraphs)')
B('§1 Adolescent anxiety burden + Asian context — Xu 2021, Wen 2020, '
  'Saikia 2023 (verified PDF), Bhardwaj 2020')
B('§2 Vietnam gap: limited item-level + normative data — V-NAMHS 2022, '
  'Hoang Trung Hoc 2025')
B('§3 Present study: 3 RQs (item distribution, demographic differences, '
  'screening targets)')

H3('Methods (~1.000 từ)')
B('2.1 Participants: 1.352 HS THCS Hà Nội (giống Q1 — same dataset, '
  'cite Q1 paper if Q1 đã đăng trước)')
B('2.2 Instrument: RCADS Vietnamese (Chorpita 2000); 15 items in 3 subtypes; '
  '4-point Likert → 0-100')
B('2.3 Reliability: Cronbach α + McDonald ω per subtype')
B('2.4 Analytic strategy:', 0)
B('Descriptive: M, SD, range, item ranking', 1)
B('Gender: independent t-test', 1)
B('Grade: one-way ANOVA + post-hoc Tukey', 1)
B('Effect sizes: Cohen d (gender), η² (grade)', 1)
B('2.5 Ethics: identical to Q1 (IRB HNUE)')

H3('Results (~1.200 từ) — TẤT CẢ VERIFIED VS THẦY GỬI')
B('3.1 Sample characteristics (brief)')
B('3.2 Reliability per subtype')
B('3.3 Generalized anxiety items (Bảng 1 thầy gửi):', 0)
B('Top: RCADS4 "lo lắng khi nghĩ mình đã không làm tốt..." M=64,28, SD=29,69', 1)
B('RCADS13 "lo lắng điều gì đó tồi tệ sẽ xảy ra" M=59,62, SD=35,86', 1)
B('Bottom: RCADS35 "lo lắng về những gì sẽ xảy ra" M=45,86, SD=33,83', 1)
B('Subtype mean: M=55,82', 1)
B('3.4 Separation anxiety items (Bảng 2):', 0)
B('Top: RCADS46 "sẽ cảm thấy sợ nếu phải ở xa nhà qua đêm" M=27,88', 1)
B('RCADS17 "sợ nếu phải ngủ một mình" M=25,49', 1)
B('RCADS5 "sợ khi ở một mình ở nhà" M=25,35', 1)
B('Bottom: RCADS45 "lo lắng khi đi ngủ ban đêm" M=21,52', 1)
B('Subtype mean: M=25,06 (lowest of 3 subtypes)', 1)
B('3.5 Social anxiety items (Bảng 3):', 0)
B('Top: RCADS32 "lo lắng người khác nghĩ gì về mình" M=56,98', 1)
B('RCADS43 "sợ rằng mình sẽ làm trò cười..." M=49,26', 1)
B('Bottom: RCADS20 "lo rằng mình trông ngốc nghếch" M=42,09', 1)
B('Subtype mean: M=48,41', 1)
B('3.6 Demographic comparisons (Bảng 4):', 0)
B('Gender: female > male cho GAD (59,47 vs 51,43; F=44,484; p<0,001) + '
  'Social (52,74 vs 43,20; F=45,984; p<0,001); **separation NS** '
  '(F=0,246; p=0,620)', 1)
B('Grade trajectory: GAD tăng 6→9 (54,32 → 53,65 → 55,63 → 59,79); '
  'separation GIẢM 6→9 (32,13 → 27,14 → 20,88 → 19,46); '
  'social: F=4,879; p=0,002', 1)

H3('Discussion (~1.000 từ — 4 paragraphs)')
B('4.1 Summary: academic worry dominant GAD; social judgment dominant SocAD')
B('4.2 Developmental pattern: separation declines (DSM-5 childhood-onset), GAD '
  'increases (late-childhood/adolescence onset)')
B('4.3 Gender: female > male except separation — brief mention, reserve depth '
  'for Q1')
B('4.4 Implications: screening + limitations + future replication')

H3('References (~25-35 refs)')
B('Empirical: Xu 2021, Wen 2020, Anderson 2025, V-NAMHS 2022, Hoang Trung Hoc 2025, '
  'Bhardwaj 2020, Nakie 2022, Saikia 2023 (verified PDF)')
B('Methodological: Chorpita 2000 (RCADS), Rosenberg 1965')

H3('Tables + Figures plan')
B('Table 1: Sample demographics')
B('Table 2-4: Item-level cho 3 subtypes')
B('Table 5: Demographic + grade comparisons')
B('Figure 1: Grade trajectory plot 3 subtypes (key visual)')


d.add_page_break()


# ============================================================
# PART 3: DATA SPLIT TABLE
# ============================================================
H1('PHẦN 3 — BẢNG PHÂN CHIA DỮ LIỆU CHỐNG TRÙNG LẶP')

H2('3.1 Nguyên tắc anti-overlap (4 tiêu chí ortogonal)')
B('(a) Câu hỏi NC khác — Q1 hỏi cơ chế; Q3 hỏi mô tả')
B('(b) Statistical lens khác — Q1 SEM (β); Q3 descriptive (M, ranking, ANOVA)')
B('(c) Outcome khác — Q1 mô hình + invariance; Q3 item-level + demographic')
B('(d) Method paradigm khác — Q1 mixed-methods; Q3 quantitative descriptive')

H2('3.2 Bảng phân bổ dữ liệu')

t = d.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Dữ liệu (từ "Thực trạng RLLA" thầy gửi)'
hdr[1].text = 'Bài Q1 (BMC Psychiatry)'
hdr[2].text = 'Bài Q3 (PLOS ONE)'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

data_split = [
    ('Bảng 1 (RLLA tổng quát item-level)', 'KHÔNG dùng — chỉ summary M=55,82',
     '✓ CHỦ LỰC — full item table 7 mục'),
    ('Bảng 2 (Separation anxiety item-level)', 'KHÔNG dùng — chỉ summary M=25,06',
     '✓ CHỦ LỰC — full 4 mục'),
    ('Bảng 3 (Social anxiety item-level)', 'KHÔNG dùng — chỉ summary M=48,41',
     '✓ CHỦ LỰC — full 4 mục'),
    ('Bảng 4 (Demographic + Grade)', '✓ Background only; focus multi-group invariance',
     '✓ Full table — primary analysis'),
    ('Bảng 5 (3 YTNC item-level)', '✓ Aggregate → SEM predictors',
     'KHÔNG dùng (out of scope)'),
    ('Bảng 6 (4 YTBV item-level)', '✓ Aggregate → SEM predictors',
     'KHÔNG dùng (out of scope)'),
    ('SEM β coefficients (LA Bảng 27-42)', '✓ PRIMARY FINDING',
     'KHÔNG dùng'),
    ('Multi-group gender invariance', '✓ NOVEL FINDING (separation)',
     'Brief mention only, KHÔNG technical detail'),
    ('Qualitative interviews (NND, MNA, NCS confirm thêm)', '✓ Mixed-methods integration',
     'KHÔNG dùng'),
    ('Khung CT 8 nội dung (LA Kết luận)', '✓ Discussion implications',
     'Brief mention only'),
]
for row_data in data_split:
    row = t.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if '✓' in txt and 'KHÔNG' not in txt:
                    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                elif 'KHÔNG' in txt:
                    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


d.add_page_break()


# ============================================================
# PART 4: FACT VERIFICATION REPORT
# ============================================================
H1('PHẦN 4 — FACT VERIFICATION REPORT')

P('Mỗi fact trong outline đã được verify khớp với 2 nguồn gốc: '
  '(1) file "Thực trạng RLLA" thầy gửi 01/06/2026; '
  '(2) LA chính `01_LuanAn_v4_ChuanTrinhBay_28052026.docx`.', italic=True)

H2('4.1 Số liệu mẫu (verified vs Thực trạng + LA)')
V('1.352 HS — tổng (Thực trạng Bảng 4 + LA verified)')
V('614 HS nam — Thực trạng Bảng 4 R1')
V('738 HS nữ — Thực trạng Bảng 4 R2')
V('Khối 6: 368 / Khối 7: 316 / Khối 8: 340 / Khối 9: 328 — Thực trạng Bảng 4 R4-R7')

H2('4.2 Means + SD RLLA subtypes (verified vs Thực trạng Bảng 1+2+3)')
V('GAD: M=55,82 (SD~22,39) — tổng Bảng 1')
V('SAD (chia ly): M=25,06 (SD=24,293) — tổng Bảng 2 R5')
V('SocAD (xã hội): M=48,41 (SD=26,191) — tổng Bảng 3 R5')

H2('4.3 Item-level top/bottom (verified Bảng 1+2+3)')
V('GAD top: RCADS4 M=64,28 SD=29,69')
V('GAD 2nd: RCADS13 M=59,62 SD=35,86')
V('GAD bottom: RCADS35 M=45,86 SD=33,83')
V('SAD top: RCADS46 M=27,88')
V('SAD 2nd: RCADS17 M=25,49')
V('SAD 3rd: RCADS5 M=25,35')
V('SAD bottom: RCADS45 M=21,52')
V('SocAD top: RCADS32 M=56,98')
V('SocAD 2nd: RCADS43 M=49,26')
V('SocAD 3rd: RCADS38 M=45,32')
V('SocAD bottom: RCADS20 M=42,09')

H2('4.4 Demographic statistics (Bảng 4)')
V('Gender M (Nam): GAD M=51,43 / SAD M=25,42 / SocAD M=43,20 / Total M=40,02')
V('Gender F (Nữ): GAD M=59,47 / SAD M=24,76 / SocAD M=52,74 / Total M=45,66')
V('Gender F-tests: GAD F=44,484; SocAD F=45,984; Total F=29,642 (all p<0,001)')
V('**Separation gender NS: F=0,246; p=0,620** ← KEY NOVEL FINDING')
V('Grade trajectory GAD: 54,32 → 53,65 → 55,63 → 59,79')
V('Grade trajectory SAD: 32,13 → 27,14 → 20,88 → 19,46 (giảm dần — '
  'novel finding cho Q3)')

H2('4.5 SEM β coefficients (verified vs LA Bảng 27-42)')
V('ALHT β=+0,533 (LA Bảng 27)')
V('NĐT β=+0,400 (LA Bảng 30)')
V('BNHĐ β=+0,276 (LA Bảng 33)')
V('TTr β=-0,457 (LA Bảng 32)')
V('HTCM β=-0,273 (LA Bảng 34)')
V('GBTH β=-0,108 (LA Bảng 30)')
V('HTBB β=-0,015 (LA Bảng 36, ns)')
V('Bullying → SAD β=+0,376 (subtype-specific)')
V('R² total anxiety risk = 0,284')
V('R² total anxiety protective = 0,209')

H2('4.6 Yếu tố nguy cơ + bảo vệ item-level (Bảng 5+6)')
V('ALHT M=51,13 / ESSA.4 "việc học..." M=58,56 (top)')
V('NĐT M=28,38')
V('BNHĐ thể chất M=13,52 / Lời nói M=18,51')
V('HTCM M=57,65')
V('TTr M=54,85')
V('GBTH M=52,60')
V('PSSM.10 "tham gia hoạt động" M=57,69 (top)')
V('PSSM.8 "thân thiện" M=57,40')
V('PSSM.1 "thuộc về" M=47,86 (low)')
V('PSSM.5 "GV quan tâm" M=47,36 (low)')

H2('4.7 17 Citations đều có PDF verified')
V('Xu et al. 2021 — QT010_Xu_2021_China_LargestEpi.pdf')
V('Chen et al. 2023 — QT007_Chen_et_al_2023_China_BMCPsych.pdf')
V('Wen et al. 2020 — QT008_Wen_2020_China_Rural_LPA.pdf')
V('Anderson et al. 2025 — QT014_Anderson_2025_Wiley_Narrative.pdf')
V('Compas et al. 2017 — Compas_2017_Coping_MetaAnalysis.pdf')
V('Pascoe et al. 2020 — Pascoe_2020.pdf + QT067')
V('Steare et al. 2023 — Steare_2023_AcademicPressure_SR.pdf')
V('Jefferies & Ungar 2020 — QT035_Jefferies_SocialAnxiety_7Countries_2020.pdf')
V('Nakie et al. 2022 — QT006_Nakie_et_al_2022_Ethiopia.pdf')
V('V-NAMHS 2022 — V-NAMHS_2022.pdf')
V('Hoàng Trung Học 2025 — VN014_HoangTrungHoc_2025_VN_COVID.pdf')
V('Bhardwaj et al. 2020 — QT011_Bhardwaj_et_al_2020_India_Chandigarh.pdf')
V('Qiu et al. 2022 — QT009_Qiu_et_al_2022_China_Parenting.pdf')
V('Zhu 2025 — QT015_Zhu_2025_BMC_China.pdf')
V('Alharbi et al. 2019 — 05_Alharbi_et_al_2019.pdf')
V('**Saikia et al. 2023 — 11_Saikia_2023_IJCM.pdf** (REAL Saikia, không phải '
  'file QT002 mislabeled)')
V('GBD ASEAN 2025 — DICH_GBD_2021_ASEAN_Lancet_2025.md + duplicate PDFs trong kho')


d.add_page_break()


# ============================================================
# PART 5: ETHICS HNUE EXPLAINED
# ============================================================
H1('PHẦN 5 — ETHICS APPROVAL HNUE GIẢI THÍCH')

H2('5.1 Định nghĩa')
P('**Phê duyệt Đạo đức Nghiên cứu** (Research Ethics Approval / IRB approval) '
  '— văn bản chính thức từ Hội đồng Đạo đức (Institutional Review Board) của '
  'Trường ĐHSPHN cho phép nghiên cứu trên người, đặc biệt trẻ vị thành niên.')

H2('5.2 Tại sao bắt buộc cho Q1 + Q3')
B('BMC Psychiatry + PLOS ONE đều BẮT BUỘC declaration: số quyết định IRB, '
  'ngày phê duyệt, tên hội đồng')
B('Liên quan Helsinki Declaration + Vietnamese Law on Children 2016')
B('Thiếu → reject ngay từ editor screening hoặc reviewer flag')

H2('5.3 5 documents chuẩn cần có')
B('(1) Quyết định phê duyệt từ Hội đồng Đạo đức ĐHSPHN (số/ký hiệu, ngày)')
B('(2) Bản consent form parents/legal guardians đã ký')
B('(3) Bản assent form học sinh đã ký')
B('(4) Letter từ Ban Giám hiệu 2 trường (Nhật Tân + Tây Mỗ)')
B('(5) Quy trình anonymization data')

H2('5.4 Tình huống NCS có thể đang ở')
t3 = d.add_table(rows=1, cols=2)
t3.style = 'Light Grid Accent 1'
hdr3 = t3.rows[0].cells
hdr3[0].text = 'Tình huống'
hdr3[1].text = 'Hành động cần'
for c in hdr3:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

ethics_situations = [
    ('Đã có letter IRB chính thức từ HNUE',
     'Tốt — chỉ cần copy số/ngày vào Methods + ZIP kèm submit'),
    ('Chỉ có consent + BGH approval, không IRB letter',
     'CẦN xin retroactive IRB approval từ HNUE (2-4 tuần)'),
    ('Không có gì',
     'BMC + PLOS sẽ REJECT. Phải xin IRB letter TRƯỚC khi submit'),
]
for row_data in ethics_situations:
    row = t3.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)

H2('5.5 Em đề xuất NCS làm sớm')
B('(1) Liên hệ Phòng Quản lý Khoa học HNUE hỏi quy trình retroactive IRB '
  'cho LA TS')
B('(2) Chuẩn bị: tóm tắt đề tài + consent form đã dùng + danh sách 2 trường '
  '+ LA chính')
B('(3) Submit IRB approval request — wait 2-4 tuần')
B('(4) Trong khi chờ, em viết outline + draft (KHÔNG submit cho đến khi có IRB)')


d.add_page_break()


# ============================================================
# PART 6: TIMELINE + NEXT STEPS
# ============================================================
H1('PHẦN 6 — TIMELINE + BƯỚC TIẾP THEO')

H2('6.1 Timeline đề xuất')
t4 = d.add_table(rows=1, cols=3)
t4.style = 'Light Grid Accent 1'
hdr4 = t4.rows[0].cells
hdr4[0].text = 'Tuần'; hdr4[1].text = 'Việc Q1'; hdr4[2].text = 'Việc Q3'
for c in hdr4:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

timeline = [
    ('Tuần 1 (01-07/06)',
     'Bố cục v2 đã xong; NCS xin IRB letter; em viết outline chi tiết Q1',
     '(chưa bắt đầu)'),
    ('Tuần 2 (08-14/06)',
     'Em viết Methods + Results',
     'Em viết outline Q3'),
    ('Tuần 3 (15-21/06)',
     'Em viết Discussion + Intro; NCS review',
     'Em viết Methods + Results'),
    ('Tuần 4 (22-28/06)',
     'Verify fact + plagiarism + format BMC Psychiatry',
     'Em viết Discussion + Intro'),
    ('Tuần 5 (29/06-05/07)',
     'Submit Q1 (cần IRB letter đã có)',
     'Verify + plagiarism check'),
    ('Tuần 6 (06-12/07)',
     'Respond peer review',
     'Submit Q3 (cần IRB letter)'),
]
for row_data in timeline:
    row = t4.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)


H2('6.2 Câu hỏi NCS sẽ trả lời sau (note)')
B('(6) Qualitative interviews: số lượng + transcripts có chưa? — '
  'NCS sẽ confirm sau')


H2('6.3 Cảnh báo academic integrity')
B('Tuyệt đối KHÔNG đụng bài Q2.5 KHGDVN của NCS (Bai1, Bai2)')
B('KHÔNG copy paraphrase từ LA chính — 2 bài viết hoàn toàn original tiếng Anh')
B('Plagiarism check tự kiểm: tránh chuỗi ≥7 từ giống bất kỳ paper nào đã đăng')
B('Citation: chỉ cite verified PDF + DOI verifiable')
B('Số liệu KHÔNG được lệch khỏi LA: 1.352 HS, F=0,246, ~85-89%, 8 nội dung CT')


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
