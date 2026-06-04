"""
Rebuild rag_main_index.json từ TẤT CẢ canonical entries (84 bài) trong Tom-tat-tung-bai/.
Thay thế phiên bản cũ (47 entries từ rebuild_rag_full.py hardcoded).
Output → cả heavy + light web/data/.
"""
import sys, io, json, re, math
from pathlib import Path
from collections import Counter
from docx import Document
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
TT_DIR = BASE / 'Tom-tat-tung-bai'
WEB_HEAVY = BASE / 'tro-ly-nghien-cuu-tam-ly' / 'web' / 'data'
WEB_LIGHT = BASE / 'tro-ly-nghien-cuu-tam-ly-light' / 'web' / 'data'
LOCAL_DATA = BASE / '06_Scripts' / 'questions_kg_data'

STOPWORDS = set('''a an and or but in on at to from for of with by is are was were be been being
have has had do does did the this that these those it its they them their we us our you your
he she his her i me my là và hoặc nhưng trong tại đến từ cho của với bởi là các một những này đó
nó họ chúng ta chúng tôi tôi bạn anh ấy cô ấy có được bị đã đang sẽ cũng không chỉ rất đều như
ai ở khi nào tới sau trước giữa ngoài trên dưới bên cùng qua vào ra lên xuống đây đó kia gì thì mà'''.split())

def tokenize(text):
    if not text: return []
    text = text.lower()
    tokens = re.findall(r'\b[\wÀ-ỹđĐ]+\b', text, flags=re.UNICODE)
    return [t for t in tokens if t not in STOPWORDS and len(t) > 1]

def read_docx_text(fp, max_chars=4000):
    try:
        d = Document(str(fp))
        parts = []
        for p in d.paragraphs:
            if p.text.strip():
                parts.append(p.text.strip())
        for t in d.tables:
            for row in t.rows:
                row_text = ' | '.join(c.text.strip()[:80] for c in row.cells if c.text.strip())
                if row_text:
                    parts.append(row_text)
        text = '\n'.join(parts)
        return text[:max_chars]
    except Exception as e:
        return f'[ERROR reading docx: {e}]'

# === Load canonical ===
with open(IDX_PATH, encoding='utf-8') as f:
    canonical = json.load(f)

print(f'Loaded canonical_index: {len(canonical)} entries')

# === Build entries from Tom-tat ===
entries = []
missing_summary = []

for cid in sorted(canonical.keys()):
    meta = canonical[cid]
    summary_name = meta.get('summary')
    if not summary_name:
        missing_summary.append(cid)
        continue
    summary_path = TT_DIR / summary_name
    if not summary_path.exists():
        # Try fallback names
        candidates = list(TT_DIR.glob(f'{cid}*.docx'))
        if candidates:
            summary_path = candidates[0]
        else:
            missing_summary.append(cid)
            continue
    text = read_docx_text(summary_path, max_chars=4500)
    if text.startswith('[ERROR'):
        missing_summary.append(cid + ' (read error)')
        continue

    # Try parse title from text (first non-meta line)
    lines = [l.strip() for l in text.split('\n') if l.strip()]
    title = ''
    for l in lines[:8]:
        if 'Tóm tắt' in l or 'Summary' in l or l.startswith(cid):
            continue
        if 20 < len(l) < 300:
            title = l
            break
    if not title and lines:
        title = lines[0][:200]

    tokens = dict(Counter(tokenize(text)))
    entries.append({
        'id': cid,
        'type': 'paper',
        'title': title or meta.get('descriptor', ''),
        'text': text,
        'tokens': tokens,
        'meta': {
            'id': cid,
            'descriptor': meta.get('descriptor', ''),
            'pdf': meta.get('pdf'),
            'pdf_folder': meta.get('pdf_folder'),
            'summary_file': summary_name,
            'translation_file': meta.get('translation'),
            'doi': meta.get('doi'),
            'source_url': meta.get('source_url'),
            'scope': meta.get('scope', 'main'),
            'region_code': 'VN' if cid.startswith('VN') else 'QT',
        },
    })

# === Insights — pull from old rebuild_rag_full.py ===
insights_block = []
old_rebuild = BASE / 'tro-ly-nghien-cuu-tam-ly' / 'rebuild_rag_full.py'
if old_rebuild.exists():
    src = old_rebuild.read_text(encoding='utf-8')
    m = re.search(r'^insights\s*=\s*\[', src, re.MULTILINE)
    if m:
        i = m.end() - 1
        depth = 0
        while i < len(src):
            c = src[i]
            if c == '[':
                depth += 1
            elif c == ']':
                depth -= 1
                if depth == 0:
                    block = src[m.start():i+1]
                    break
            i += 1
        ns = {'read_docx': lambda *a, **kw: '', 'os': __import__('os'), 'BASE': str(old_rebuild.parent)}
        try:
            exec(block, ns)
            insights_block = ns['insights']
        except Exception as e:
            print(f'Insights parse failed: {e}')

for ins in insights_block:
    text = f"{ins.get('title', '')}\n{ins.get('content', '')}"
    tokens = dict(Counter(tokenize(text)))
    entries.append({
        'id': ins.get('id', ''),
        'type': 'insight',
        'title': ins.get('title', ''),
        'text': text,
        'tokens': tokens,
        'meta': {
            'id': ins.get('id', ''),
            'region': ins.get('region', ''),
            'region_code': ins.get('region_code', ''),
            'scope': 'insight',
        },
    })

# === Compute IDF ===
N = len(entries)
df = Counter()
for e in entries:
    for tok in e['tokens']:
        df[tok] += 1
idf = {tok: math.log((N + 1) / (dfi + 1)) + 1 for tok, dfi in df.items()}

n_papers = sum(1 for e in entries if e['type'] == 'paper')
n_insights = sum(1 for e in entries if e['type'] == 'insight')

out_data = {
    'meta': {
        'created': '2026-04-24',
        'version': 'v2-canonical-84',
        'n_entries': N,
        'n_papers': n_papers,
        'n_insights': n_insights,
        'total_unique_tokens': len(idf),
        'description': f'TF-IDF index của {n_papers} canonical papers (24 VN + 60 QT) + {n_insights} insights',
    },
    'idf': idf,
    'entries': entries,
}

# Save
for dest in [LOCAL_DATA / 'rag_main_index.json',
             WEB_HEAVY / 'rag_main_index.json',
             WEB_LIGHT / 'rag_main_index.json']:
    if dest.parent.exists():
        with open(dest, 'w', encoding='utf-8') as f:
            json.dump(out_data, f, ensure_ascii=False, indent=1)
        size = dest.stat().st_size / 1024
        print(f'Saved: {dest} ({size:.1f} KB)')

print(f'\n=== STATS ===')
print(f'Total entries: {N}')
print(f'  Papers: {n_papers}')
print(f'  Insights: {n_insights}')
print(f'Unique tokens: {len(idf)}')
print(f'Missing summaries: {len(missing_summary)}')
if missing_summary:
    print(f'  → {missing_summary}')

# Smoke test
print('\n=== SMOKE TEST ===')
for q in ['CBT lo âu vị thành niên', 'GBD ASEAN xu hướng', 'DASS-21 Việt Nam',
          'NNT can thiệp', 'mindfulness resilience', 'gender post pandemic']:
    qt = tokenize(q)
    scored = []
    for e in entries:
        s = sum(e['tokens'].get(t, 0) * idf.get(t, 1.0) for t in qt)
        if s > 0:
            scored.append((s, e))
    scored.sort(reverse=True, key=lambda x: x[0])
    print(f"\nQ: '{q}'")
    for s, e in scored[:3]:
        print(f'  [{s:.1f}] {e["id"]}: {e["title"][:70]}')
