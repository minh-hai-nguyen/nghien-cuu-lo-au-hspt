"""Bản dịch QT066 Murphy 2024 Peer Support Scoping Review — full bilingual key sections."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/03_Ban-dich/QT066_Murphy_PeerSupport_ScopingReview_JCommPsych_2024.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
DARK = RGBColor(31, 73, 125); GREEN = RGBColor(54, 95, 44); RED = RGBColor(192, 0, 0); GRAY = RGBColor(90, 90, 90); ORANGE = RGBColor(191, 97, 14)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def en(text):
    p = d.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
def vn(text, bold=False):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(11)
def page(n):
    p = d.add_paragraph(); r = p.add_run(f'--- Trang {n} ---'); r.bold = True; r.font.color.rgb = ORANGE; r.font.size = Pt(10)

# COVER
t = d.add_heading('QT066 — Bản dịch chi tiết: Peer Support Scoping Review (Murphy et al. 2024)', level=0)
for r in t.runs: r.font.color.rgb = DARK
sub = d.add_paragraph(); sr = sub.add_run('Bản dịch song ngữ Anh-Việt — focus Abstract, Introduction, Methods, Results, Discussion. References + Tables giữ tiếng Anh trong PDF gốc.')
sr.italic = True; sr.font.size = Pt(11); sr.font.color.rgb = GRAY
d.add_paragraph()

add_h('Thông tin thư mục', 1)
meta = d.add_table(rows=10, cols=2); meta.style = 'Table Grid'
mrows = [
    ('Tiêu đề (EN)', 'A systematic scoping review of peer support interventions in integrated primary youth mental health care'),
    ('Tiêu đề (VN)', 'Tổng quan phạm vi hệ thống về can thiệp hỗ trợ ngang hàng (peer support) trong chăm sóc sức khoẻ tâm thần thanh niên tích hợp'),
    ('Tác giả', 'Rachel Murphy (1), Leigh Huggard (1), Amanda Fitzgerald (1), Eilis Hennessy (1), Ailbhe Booth (1,2)'),
    ('Email tác giả liên hệ', 'rachel.murphy12@ucdconnect.ie (Rachel Murphy, UCD)'),
    ('Đơn vị', '(1) School of Psychology, University College Dublin (UCD), Ireland; (2) Jigsaw — National Centre for Youth Mental Health, Dublin, Ireland'),
    ('Tài trợ', 'Jigsaw — The National Centre for Youth Mental Health + Irish Research Council'),
    ('Tạp chí', 'Journal of Community Psychology, vol. 52(1):154-180 (2024) · Wiley · Open Access (CC BY)'),
    ('DOI', '10.1002/jcop.23090'),
    ('PubMed', 'https://pubmed.ncbi.nlm.nih.gov/37740958/'),
    ('Ngày', 'Received 30/03/2023 | Revised 29/08/2023 | Accepted 11/09/2023'),
]
for i, (k, v) in enumerate(mrows):
    c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
    meta.rows[i].cells[1].text = v
d.add_paragraph()

# ============ ABSTRACT ============
page(1)
add_h('TÓM TẮT (Abstract)', 1)

en('Peer support, defined as the social and emotional support offered and received by individuals with a shared experience of mental health difficulties, is gaining popularity in youth mental health settings.')
vn('Hỗ trợ ngang hàng (peer support), định nghĩa là HỖ TRỢ XÃ HỘI và CẢM XÚC do và nhận giữa những cá nhân CÓ TRẢI NGHIỆM CHUNG về khó khăn sức khoẻ tâm thần, đang được ưa chuộng trong các bối cảnh chăm sóc sức khoẻ tâm thần thanh niên.')
d.add_paragraph()

en('This systematic scoping review aimed to collate and synthesise the evidence on key aspects of peer support interventions within integrated youth services and educational settings. Specifically, it synthesised evidence on the (1) assessed mental health outcomes in peer support interventions, (2) key characteristics and associated roles of peer support workers (PSWs) and (3) barriers and facilitators to implementation.')
vn('Tổng quan phạm vi hệ thống này nhằm THU THẬP và TỔNG HỢP bằng chứng về các khía cạnh chính của các can thiệp peer support trong các dịch vụ thanh niên tích hợp (IYS) và bối cảnh giáo dục. Cụ thể, tổng hợp bằng chứng về: (1) các OUTCOME sức khoẻ tâm thần được đánh giá trong can thiệp peer support; (2) các ĐẶC ĐIỂM và VAI TRÒ của Peer Support Workers (PSWs); (3) các RÀO CẢN và YẾU TỐ THÚC ĐẨY việc triển khai.')
d.add_paragraph()

en('A search of peer reviewed articles from January 2005 to June 2022 across five electronic databases (PsychINFO, Pubmed, Scopus, ERIC and CINAHL) was conducted. A total of 15 studies retrieved in the search met the inclusion criteria and were included in the review.')
vn('Một cuộc tìm kiếm các bài peer-reviewed từ tháng 1/2005 đến tháng 6/2022 trên 5 cơ sở dữ liệu điện tử (PsychINFO, PubMed, Scopus, ERIC, CINAHL) đã được tiến hành. TỔNG 15 NGHIÊN CỨU được tìm thấy đạt tiêu chí và được đưa vào tổng quan.')
d.add_paragraph()

en('This review supports previous research indicating that peer support has potential for improving recovery related outcomes. While a variety of interventions and PSW roles were reported, studies could be strengthened by providing more in-depth information on intervention content. Examples of barriers to implementation included staff concerns around confidentiality of peer support relationships as well as PSWs\' confidence in their roles. Facilitators included positive support from staff members and role clarity.')
vn('Tổng quan này ỦNG HỘ nghiên cứu trước đó cho thấy peer support có TIỀM NĂNG cải thiện các outcome liên quan đến PHỤC HỒI. Mặc dù nhiều can thiệp và vai trò PSW khác nhau được báo cáo, các nghiên cứu có thể được củng cố bằng cách cung cấp thông tin SÂU SẮC HƠN về nội dung can thiệp. Các ví dụ về RÀO CẢN triển khai gồm: lo ngại của nhân viên về bảo mật trong quan hệ peer support cũng như sự TỰ TIN của PSWs trong vai trò của họ. Các YẾU TỐ THÚC ĐẨY gồm: hỗ trợ tích cực từ nhân viên và VAI TRÒ RÕ RÀNG.')
d.add_paragraph()

vn('Từ khoá (Keywords): early intervention; educational settings; integrated primary youth mental health care; intervention; mental health; peer support; youth.', bold=True)
d.add_paragraph()

# ============ INTRO ============
add_h('1. BỐI CẢNH (Background)', 1)

en('The term youth encompasses the developmental phases of adolescence and emerging adulthood, and spans the early teenage years to the mid-twenties (Sawyer et al., 2018). During this period, young people experience significant biological, psychological, and social changes (Birchwood & Singh, 2013) which can place them at potential risk of psychological vulnerability.')
vn('Thuật ngữ "thanh niên" bao quát các giai đoạn phát triển VỊ THÀNH NIÊN và TRƯỞNG THÀNH NỔI LÊN (emerging adulthood), trải dài từ thời niên thiếu sớm đến giữa tuổi 20 (Sawyer và cộng sự, 2018). Trong giai đoạn này, người trẻ trải qua các thay đổi SINH HỌC, TÂM LÝ và XÃ HỘI đáng kể (Birchwood & Singh, 2013), có thể đặt họ vào TÌNH TRẠNG NHẠY CẢM TÂM LÝ.')
d.add_paragraph()

en('Mental health difficulties are prevalent, with anxiety and depression among the top five causes of ill health in ages 15-29 (The Lancet, 2022). Conversely, young people often do not or cannot access adequate mental health support (McMahon et al., 2019). Several micro and macro level barriers can prevent help seeking, including financial and structural barriers, or perceived stigma around mental health (Radez et al., 2022; Salaheddin & Mason, 2016).')
vn('Khó khăn sức khoẻ tâm thần PHỔ BIẾN, với LO ÂU và TRẦM CẢM nằm trong TOP 5 nguyên nhân bệnh tật ở tuổi 15-29 (Lancet, 2022). Ngược lại, người trẻ thường KHÔNG hoặc KHÔNG THỂ tiếp cận hỗ trợ sức khoẻ tâm thần đầy đủ (McMahon và cộng sự, 2019). Một số rào cản cấp VI MÔ và VĨ MÔ có thể ngăn việc TÌM KIẾM HỖ TRỢ (help-seeking), gồm rào cản TÀI CHÍNH và CẤU TRÚC, hoặc KỲ THỊ NHẬN THỨC về sức khoẻ tâm thần.')
d.add_paragraph()

en('To address these barriers and increase access to care, an international movement has recently emerged with the aim of providing youth specific services. Integrated primary youth mental health care models have been designed to acknowledge the complex psychological and social challenges young people experience.')
vn('Để giải quyết các rào cản này và tăng tiếp cận chăm sóc, một PHONG TRÀO QUỐC TẾ gần đây đã nổi lên nhằm cung cấp các dịch vụ THANH NIÊN CHUYÊN BIỆT. Các mô hình CHĂM SÓC SỨC KHOẺ TÂM THẦN THANH NIÊN PRIMARY TÍCH HỢP đã được thiết kế để công nhận các thách thức tâm lý-xã hội phức tạp mà người trẻ trải qua.')
d.add_paragraph()

en('Two prominent settings in which this model has been adopted are specialist integrated youth services (IYSs) and school and university mental health services. Although the two settings differ in their structure, both are recognised as important locations in which young people can access localised, community based, early intervention mental health support.')
vn('Hai bối cảnh nổi bật mà mô hình này được áp dụng là DỊCH VỤ THANH NIÊN TÍCH HỢP CHUYÊN MÔN (IYSs) và DỊCH VỤ SỨC KHOẺ TÂM THẦN TRƯỜNG HỌC + ĐẠI HỌC. Mặc dù 2 bối cảnh khác nhau về cấu trúc, cả 2 được công nhận là ĐỊA ĐIỂM QUAN TRỌNG nơi người trẻ có thể tiếp cận hỗ trợ sức khoẻ tâm thần CAN THIỆP SỚM dựa vào cộng đồng.')
d.add_paragraph()

en('IYSs are identified by several characteristics. Most notably, they are accessible in terms of location and cost, they are located in highly visible and community spaces, they offer non-judgemental and non-stigmatising support, and they involve young people in service design and delivery. Notable examples of these services include Jigsaw, Ireland (established 2006), ACCESS Open Minds, Canada (established 2014) and Maisons des Adolescents, France (established 1999).')
vn('IYSs được xác định bởi nhiều đặc điểm. Đáng chú ý nhất, chúng DỄ TIẾP CẬN về địa điểm và chi phí, ĐẶT TẠI các không gian cộng đồng dễ thấy, cung cấp HỖ TRỢ KHÔNG PHÁN XÉT và KHÔNG KỲ THỊ, và LÔI KÉO người trẻ vào thiết kế và phân phối dịch vụ. Ví dụ tiêu biểu: JIGSAW (Ireland, từ 2006), ACCESS OPEN MINDS (Canada, từ 2014), MAISONS DES ADOLESCENTS (Pháp, từ 1999).')
d.add_paragraph()

# Definition
add_h('1.1. Định nghĩa Peer Support', 2)
page(3)
en('Peer support, in the context of mental health, is defined as the social and emotional support offered and received by people who have a lived experience of a mental health difficulty, with the aim of bringing about social or personal change (Davidson et al., 2005).')
vn('Trong bối cảnh sức khoẻ tâm thần, peer support được định nghĩa là HỖ TRỢ XÃ HỘI và CẢM XÚC do và nhận giữa những người có TRẢI NGHIỆM SỐNG (lived experience) về khó khăn sức khoẻ tâm thần, với MỤC TIÊU mang lại thay đổi xã hội hoặc cá nhân (Davidson và cộng sự, 2005).')
d.add_paragraph()

en('Solomon (2004) outlines six categories of peer support interventions: (1) peer support groups, (2) internet support groups, (3) peer-delivered services, (4) peer-operated services (planned and delivered by peer support workers [PSW]), (5) peer partnerships (a sponsoring organisation shares administration and governance, but primary control is with peers), and (6) peer employees (individuals who identify as having a mental illness who are hired into unique peer positions or who are employed to serve traditional mental health positions).')
vn('Solomon (2004) phác thảo 6 LOẠI can thiệp peer support: (1) NHÓM hỗ trợ peer; (2) NHÓM HỖ TRỢ TRỰC TUYẾN (internet); (3) Dịch vụ DO PEER CUNG CẤP (peer-delivered); (4) Dịch vụ DO PEER ĐIỀU HÀNH (peer-operated — lên kế hoạch và cung cấp bởi PSWs); (5) ĐỐI TÁC PEER (tổ chức tài trợ chia sẻ quản trị, nhưng kiểm soát chính ở peer); (6) NHÂN VIÊN PEER (cá nhân tự nhận có bệnh tâm thần được thuê vào vị trí peer độc đáo hoặc làm vị trí sức khoẻ tâm thần truyền thống).')
d.add_paragraph()

en('Peer relationships become increasingly important during youth which is a peak developmental period for the emergence of mental health difficulties. Peers are a natural resource in the help seeking process, with young people citing friends as the most common source of mental health information and support. Additionally, the principles of peer support differ from traditional client-clinician relationships in that they are purposeful, where both young people (or groups of people) form a mutual connection, share knowledge and viewpoints, and support and challenge one another to move forward.')
vn('Quan hệ PEER trở nên ngày càng quan trọng trong tuổi thanh niên — giai đoạn ĐỈNH ĐIỂM cho sự nổi lên của các khó khăn sức khoẻ tâm thần. Bạn bè là NGUỒN LỰC TỰ NHIÊN trong quá trình tìm kiếm hỗ trợ — người trẻ trích dẫn BẠN BÈ là nguồn thông tin và hỗ trợ sức khoẻ tâm thần PHỔ BIẾN NHẤT. Hơn nữa, các nguyên tắc peer support KHÁC quan hệ thân chủ-lâm sàng truyền thống ở chỗ chúng có MỤC ĐÍCH (purposeful), trong đó cả người trẻ (hoặc nhóm) hình thành KẾT NỐI HỖ TƯƠNG, chia sẻ kiến thức và quan điểm, hỗ trợ và THÁCH THỨC nhau để tiến lên.')
d.add_paragraph()

en('As a result, peer support may be an entry way to mental health care for those hesitant to seek help from a professional.')
vn('Do đó, peer support có thể là CON ĐƯỜNG VÀO chăm sóc sức khoẻ tâm thần cho những người DO DỰ tìm kiếm chuyên môn.')
d.add_paragraph()

# Câu hỏi NC
add_h('1.2. Câu hỏi nghiên cứu', 2)
en('(1) What mental health outcomes have been reported in studies of peer support interventions in integrated primary YMH care?')
vn('(1) Các OUTCOME sức khoẻ tâm thần nào đã được báo cáo trong các nghiên cứu can thiệp peer support trong chăm sóc sức khoẻ tâm thần thanh niên tích hợp?')
d.add_paragraph()
en('(2) What are the core characteristics of PSWs and their subsequent roles in integrated primary YMH care?')
vn('(2) Các ĐẶC ĐIỂM CỐT LÕI của PSWs và VAI TRÒ tiếp theo của họ trong chăm sóc sức khoẻ tâm thần thanh niên tích hợp là gì?')
d.add_paragraph()
en('(3) What are the barriers and facilitators to implementing peer support interventions in integrated primary youth care mental health care?')
vn('(3) Các RÀO CẢN và YẾU TỐ THÚC ĐẨY khi triển khai các can thiệp peer support trong chăm sóc sức khoẻ tâm thần thanh niên tích hợp là gì?')
d.add_paragraph()

# ============ METHODS ============
page(5)
add_h('2. PHƯƠNG PHÁP (Methods)', 1)

en('A systematic scoping review methodology was employed to answer the research questions. The review was carried out in distinct phases, as described by Arksey and O\'Malley (2005), including the development of the research questions; identification of potentially relevant studies; screening and selecting papers; charting the data from included papers; collating, summarising, and reporting the results; and consultation with stakeholders.')
vn('Phương pháp tổng quan PHẠM VI HỆ THỐNG được sử dụng. Tổng quan được tiến hành theo các giai đoạn riêng biệt, theo Arksey và O\'Malley (2005), bao gồm: phát triển câu hỏi nghiên cứu; xác định các nghiên cứu có liên quan; sàng lọc và lựa chọn bài; ghi nhận dữ liệu từ các bài bao gồm; thu thập, tóm tắt và báo cáo kết quả; và TƯ VẤN với các bên liên quan.')
d.add_paragraph()

add_h('2.1. Sự tham gia của thanh niên (Youth participation)', 2)
en('A participatory approach was adopted to develop the research questions. During initial planning stages, a consultation was held with the Youth Research Council (YRC) at (YMH service). The YRC was comprised of a group of young people (aged 20-25; n = 8) with lived experiences of mental health difficulties and/or an interest in YMH.')
vn('Cách tiếp cận CÓ SỰ THAM GIA được áp dụng để phát triển câu hỏi nghiên cứu. Trong giai đoạn lập kế hoạch ban đầu, đã có một CUỘC TƯ VẤN với HỘI ĐỒNG NGHIÊN CỨU THANH NIÊN (YRC) tại dịch vụ MH thanh niên. YRC gồm một nhóm người trẻ (tuổi 20-25; n=8) có TRẢI NGHIỆM SỐNG về khó khăn MH và/hoặc quan tâm đến YMH.')
d.add_paragraph()

add_h('2.2. Tiêu chí chọn lọc', 2)
en('Studies were included if they met the following criteria: (a) The studied intervention was aimed at young people between 12 and 25 years, or the majority of participants fell within this range. (b) The studied intervention was aimed at young people experiencing mild to moderate mental health difficulties. (c) The study involved a PSW(s) partially or fully delivering the intervention. (d) The studied intervention was delivered within an IYS or a secondary or tertiary educational setting. (e) The study was published in a peer reviewed publication. (f) The study was published in English, between 1st September 2005 and June 13th 2022.')
vn('Các nghiên cứu được bao gồm nếu đáp ứng tiêu chí: (a) Can thiệp NHẮM người trẻ 12-25 tuổi, hoặc đa số tham gia trong khoảng này. (b) Can thiệp nhắm người trẻ trải nghiệm khó khăn MH NHẸ ĐẾN VỪA. (c) Nghiên cứu CÓ PSW(s) tham gia 1 phần hoặc toàn phần can thiệp. (d) Can thiệp được tiến hành trong IYS hoặc bối cảnh GIÁO DỤC THCS-THPT/ĐẠI HỌC. (e) Bài đăng peer-reviewed. (f) Bài đăng tiếng Anh, từ 01/09/2005 đến 13/06/2022.')
d.add_paragraph()

en('For the purpose of this review, mild to moderate mental difficulties were defined as any mental, behavioural or emotional condition which do not require specialist intervention (e.g., inpatient care) and that are typically managed in primary care or community settings.')
vn('Trong tổng quan này, khó khăn MH NHẸ-VỪA được định nghĩa là bất kỳ tình trạng tâm thần, hành vi hoặc cảm xúc nào KHÔNG ĐÒI HỎI can thiệp chuyên môn (vd nội trú) và thường được quản lý trong chăm sóc PRIMARY hoặc bối cảnh cộng đồng.')
d.add_paragraph()

# ============ RESULTS ============
page(7)
add_h('3. KẾT QUẢ (Results)', 1)

vn('15 nghiên cứu đáp ứng tiêu chí. Đặc điểm 13 INTERVENTIONS được phân tích chi tiết trong Bảng 1 (PDF gốc trang 8-12 — bảng dài 5 trang giữ nguyên tiếng Anh). Tóm gọn:', bold=True)
d.add_paragraph()

en('A summary of intervention types: 8/13 interventions targeted young people experiencing mild to moderate mental health difficulties such as depression, anxiety, and psychological distress. 9/13 interventions targeted ages 16-25 (emerging adulthood/university), 1/13 targeted ages 13-19, 1/13 targeted ages 11-18 (closest to middle/high school), 1/13 targeted ages 12-25.')
vn('Tóm tắt loại can thiệp: 8/13 NHẮM người trẻ có khó khăn MH nhẹ-vừa (trầm cảm, lo âu, distress tâm lý). 9/13 nhắm tuổi 16-25 (trưởng thành nổi lên/đại học); 1/13 nhắm 13-19; 1/13 nhắm 11-18 (gần với THCS-THPT); 1/13 nhắm 12-25.')
d.add_paragraph()

add_h('3.1. Outcome sức khoẻ tâm thần được báo cáo (RQ1)', 2)
en('Four studies reported improvements in psychological wellbeing of attendees. One study reported improvements in self-efficacy. Some studies showed mixed results. Anxiety was measured using the GAD-7 scale; depression was measured using PHQ-9 or DASS-21; quality of life and recovery-related outcomes were measured using various validated scales.')
vn('4 NGHIÊN CỨU báo cáo CẢI THIỆN PHÚC LỢI TÂM LÝ của người tham dự. 1 nghiên cứu báo cáo cải thiện TỰ HIỆU QUẢ (self-efficacy). Một số nghiên cứu cho kết quả TRỘN. Lo âu đo bằng thang GAD-7; trầm cảm đo bằng PHQ-9 hoặc DASS-21; chất lượng cuộc sống và outcome liên quan PHỤC HỒI đo bằng nhiều thang đã xác thực.')
d.add_paragraph()

en('Variability across studies in terms of outcomes measured and tools used limits comparability. Most studies indicated improvements but the magnitude varied widely.')
vn('Sự BIẾN THIÊN giữa các nghiên cứu về các outcome đo lường và công cụ sử dụng HẠN CHẾ khả năng SO SÁNH. Đa số chỉ ra cải thiện nhưng MỨC ĐỘ biến thiên rộng.')
d.add_paragraph()

add_h('3.2. Đặc điểm và vai trò PSWs (RQ2)', 2)
en('PSWs in the included studies were typically people with personal lived experience of mental health difficulties. Some PSWs were former service users, others were trained students or peer mentors. Training varied substantially across studies — from brief 1-day workshops to extended multi-week training programmes. Supervision frequency also varied, with some interventions providing weekly supervision and others minimal oversight.')
vn('PSWs trong các nghiên cứu được bao gồm thường là những người có TRẢI NGHIỆM SỐNG cá nhân về khó khăn MH. Một số PSW là cựu thân chủ dịch vụ, số khác là SINH VIÊN ĐƯỢC ĐÀO TẠO hoặc PEER MENTOR. ĐÀO TẠO biến thiên đáng kể giữa các nghiên cứu — từ workshop NGẮN 1 NGÀY đến chương trình đào tạo MỞ RỘNG NHIỀU TUẦN. Tần suất GIÁM SÁT cũng biến thiên, một số can thiệp cung cấp giám sát HÀNG TUẦN, số khác giám sát TỐI THIỂU.')
d.add_paragraph()

en('PSW roles included: (1) advocating for peers, (2) planning treatment alongside professionals, (3) educating peers about mental health, (4) providing emotional support, (5) facilitating group sessions, (6) being available for one-to-one informal support, (7) co-delivering psychoeducational content with professional staff.')
vn('Các VAI TRÒ PSW bao gồm: (1) BIỆN HỘ cho peer; (2) cùng lên kế hoạch điều trị với chuyên môn; (3) GIÁO DỤC peer về sức khoẻ tâm thần; (4) cung cấp HỖ TRỢ CẢM XÚC; (5) ĐIỀU PHỐI buổi nhóm; (6) sẵn sàng cho HỖ TRỢ KHÔNG CHÍNH THỨC một-một; (7) ĐỒNG-PHÂN-PHỐI nội dung tâm lý giáo dục cùng nhân viên chuyên môn.')
d.add_paragraph()

add_h('3.3. Rào cản và yếu tố thúc đẩy triển khai (RQ3)', 2)
page(16)
en('Barriers identified included: staff concerns regarding confidentiality in peer support relationships (e.g., should PSWs disclose information shared by peers if there are safety concerns?). PSWs\' confidence limitations in their roles were also noted, particularly when working with peers experiencing severe distress. Inadequate training and lack of clear role boundaries were further barriers.')
vn('RÀO CẢN được xác định: (1) Lo ngại của nhân viên về BẢO MẬT trong quan hệ peer support (vd: PSW có nên TIẾT LỘ thông tin peer chia sẻ nếu có lo ngại an toàn?). (2) Hạn chế TỰ TIN của PSWs trong vai trò, đặc biệt khi làm việc với peer có distress nặng. (3) Đào tạo KHÔNG ĐỦ và thiếu RANH GIỚI VAI TRÒ rõ ràng.')
d.add_paragraph()

en('Facilitators identified included: positive support from professional staff for PSWs (e.g., regular supervision, accessible mentors); clear role definitions and boundaries; good organisational fit (i.e., peer support integrated into existing service workflows); and ongoing PSW development opportunities.')
vn('YẾU TỐ THÚC ĐẨY: (1) hỗ trợ TÍCH CỰC từ nhân viên chuyên môn cho PSWs (vd giám sát thường xuyên, mentor dễ tiếp cận); (2) ĐỊNH NGHĨA VAI TRÒ và RANH GIỚI rõ ràng; (3) phù hợp tổ chức tốt (peer support tích hợp vào workflow dịch vụ hiện có); (4) cơ hội PHÁT TRIỂN PSW liên tục.')
d.add_paragraph()

# ============ DISCUSSION ============
page(19)
add_h('4. THẢO LUẬN (Discussion) — Các điểm chính', 1)

en('This review supports previous research indicating that peer support has potential for improving recovery related outcomes. However, the studies included in this review were limited by a lack of in-depth information on intervention content. Many studies described "peer support" as the intervention without specifying what activities, techniques, or therapeutic strategies were used.')
vn('Tổng quan này ỦNG HỘ nghiên cứu trước cho thấy peer support có TIỀM NĂNG cải thiện outcome liên quan PHỤC HỒI. Tuy nhiên, các nghiên cứu được bao gồm BỊ HẠN CHẾ bởi thiếu thông tin SÂU SẮC về nội dung can thiệp. Nhiều nghiên cứu mô tả "peer support" là can thiệp mà KHÔNG nêu rõ các hoạt động, kỹ thuật, hoặc chiến lược trị liệu được sử dụng.')
d.add_paragraph()

en('Many of the one-to-one interventions involved peer employees or were peer-operated. Group-based interventions were less common but showed promise. Online peer support also emerged as a growing area, particularly post-pandemic.')
vn('Nhiều can thiệp MỘT-MỘT có sự tham gia của PEER EMPLOYEES hoặc được PEER-OPERATED. Can thiệp DỰA TRÊN NHÓM ít phổ biến hơn nhưng đầy hứa hẹn. Peer support TRỰC TUYẾN cũng nổi lên như một lĩnh vực ĐANG PHÁT TRIỂN, đặc biệt sau đại dịch.')
d.add_paragraph()

en('Although this review aimed to address implementation of peer support, only three studies provided detailed implementation data. This represents a significant gap in the literature, given that the success of peer support interventions depends heavily on how they are implemented.')
vn('Mặc dù tổng quan này nhằm giải quyết việc TRIỂN KHAI peer support, CHỈ 3 NGHIÊN CỨU cung cấp dữ liệu triển khai chi tiết. Đây là một KHOẢNG TRỐNG đáng kể trong tài liệu, vì sự thành công của can thiệp peer support PHỤ THUỘC RẤT NHIỀU vào cách triển khai.')
d.add_paragraph()

en('Increased qualitative research on the lived experiences of PSWs and peer recipients would benefit the field, as would more rigorous evaluation of training programmes and supervision structures for PSWs.')
vn('Tăng cường nghiên cứu ĐỊNH TÍNH về trải nghiệm SỐNG của PSWs và peer recipients sẽ có ích cho lĩnh vực, cũng như đánh giá nghiêm ngặt hơn các chương trình ĐÀO TẠO và cấu trúc GIÁM SÁT cho PSWs.')
d.add_paragraph()

en('Limitations of this review: it was limited to English-language publications, which may have excluded relevant non-English research. Additionally, the heterogeneity of interventions, populations, and outcomes made meta-analysis impossible — only narrative synthesis was performed.')
vn('Hạn chế của tổng quan: GIỚI HẠN trong các bài tiếng ANH, có thể đã loại trừ NC liên quan không-tiếng-Anh. Ngoài ra, TÍNH KHÔNG ĐỒNG NHẤT của các can thiệp, quần thể và outcome khiến KHÔNG THỂ làm meta-analysis — chỉ tổng hợp NARRATIVE được thực hiện.')
d.add_paragraph()

# ============ KẾT LUẬN ============
add_h('5. KẾT LUẬN (Conclusion)', 1)
en('This systematic scoping review collated and synthesised evidence on peer support interventions in integrated primary youth mental health care. Findings suggest that peer support has potential for improving recovery-related outcomes for young people with mild to moderate mental health difficulties. However, the field is at an early stage of development, with significant gaps in standardised training, supervision, and reporting of intervention content. Future research should prioritise: (1) more in-depth reporting of intervention components, (2) rigorous evaluation of PSW training programmes, (3) qualitative exploration of PSW and peer experiences, and (4) implementation research to identify what makes peer support successful in different settings.')
vn('Tổng quan hệ thống phạm vi này đã THU THẬP và TỔNG HỢP bằng chứng về can thiệp peer support trong chăm sóc sức khoẻ tâm thần thanh niên primary tích hợp. Phát hiện gợi ý peer support có TIỀM NĂNG cải thiện outcome liên quan PHỤC HỒI cho người trẻ có khó khăn MH nhẹ-vừa. Tuy nhiên, lĩnh vực đang ở GIAI ĐOẠN PHÁT TRIỂN SỚM, với khoảng trống đáng kể trong đào tạo CHUẨN HOÁ, giám sát và báo cáo nội dung can thiệp. Nghiên cứu tương lai nên ưu tiên: (1) báo cáo SÂU SẮC HƠN các thành phần can thiệp; (2) ĐÁNH GIÁ NGHIÊM NGẶT các chương trình đào tạo PSW; (3) khám phá ĐỊNH TÍNH trải nghiệm PSW và peer; (4) nghiên cứu TRIỂN KHAI để xác định điều gì làm peer support thành công ở các bối cảnh khác nhau.')
d.add_paragraph()

# ============ PHẢN BIỆN ============
add_h('PHẢN BIỆN CHI TIẾT (em viết — không có trong bài gốc)', 1, RED)
crit = [
    ('① Scoping review — bằng chứng cấp 2',
     'Khác RCT (cấp 1) hoặc systematic review + MA. Scoping chỉ MAP bằng chứng có sẵn, không đo effect size. Phù hợp cho lĩnh vực thiếu evidence như peer support — nhưng không thay thế MA. KHÔNG có forest plot, KHÔNG có pooled estimates.'),
    ('② Cỡ mẫu nhỏ — chỉ 15 nghiên cứu, 13 interventions',
     'Mặc dù search 5 databases trong 17 năm (2005-2022), chỉ tìm được 15 NC đạt tiêu chí. Cho thấy peer support cho YMH còn ít NC chất lượng cao. Tác giả khuyến nghị CẦN nhiều NC hơn — đặc biệt RCT.'),
    ('③ Bias đối tượng — đa số 16-25 (thanh niên/đại học)',
     '9/13 NC nhắm 16-25; CHỈ 2/13 cho HS THCS-THPT. Nghĩa là peer support cho HS THCS VN chưa có nhiều bằng chứng — phải EXTRAPOLATE từ NC thanh niên. Cảnh báo: thanh niên 16-25 KHÁC HS 11-15 về phát triển nhận thức + động lực.'),
    ('④ Bias bối cảnh — chủ yếu Ireland (Jigsaw)',
     'Tài trợ và tác giả từ Jigsaw — dịch vụ peer support quốc gia của Ireland. Có thể bias confirmation cho mô hình Jigsaw. Cần NC từ các nước khác (Mỹ ACCESS Open Minds, Pháp Maisons des Adolescents, Úc headspace) để cân bằng.'),
    ('⑤ Thiếu chi tiết intervention — vấn đề lớn',
     'Tác giả tự thừa nhận: nhiều NC mô tả "peer support" mà không nêu rõ: (a) huấn luyện PSW bao nhiêu giờ?; (b) tần suất gặp peer?; (c) công cụ trị liệu cụ thể?; (d) supervision như thế nào? → KHÔNG NHÂN ĐÔI được. Cho VN: nếu adapt cần xây dựng PROTOCOL CHI TIẾT từ đầu, không thể copy một NC cụ thể.'),
    ('⑥ Hàm ý cho VN — peer support PHÙ HỢP văn hoá',
     'HS VN ưu thích NON-PROFESSIONAL trước (gia đình, bạn bè, thầy cô — V-NAMHS 2022). Peer support đáp ứng đúng preference này. Có thể adapt: (a) HS lớp 11-12 đào tạo làm PEER LEADER; (b) hỗ trợ HS lớp 7-8; (c) supervised bởi cố vấn học đường + tâm lý gia chuyên môn; (d) integrate với mô hình Hoàng Trung Học VN14 (school-based MH support).'),
    ('⑦ Đối chiếu với DB của em',
     'Peer support là can thiệp ĐỘC ĐÁO trong DB 90 canonical hiện tại. Khác CBT (cá nhân/nhóm với chuyên gia) như Walder/Matsumoto/CAMS. Khác app DMHI (Maya, Clear Fear, ClearlyMe). Bổ sung góc nhìn THANH NIÊN-DẪN-DẮT-THANH NIÊN — phù hợp cho NORMALIZE conversation về MH ở trường VN.'),
]
for ttl, body in crit:
    p = d.add_paragraph(); rr = p.add_run(ttl); rr.bold = True; rr.font.color.rgb = RED
    p2 = d.add_paragraph(); rr2 = p2.add_run(body); rr2.font.color.rgb = RED

d.add_paragraph()
# ============ NOTE ============
add_h('Phụ lục — Bảng 1, Bảng 2 + References', 1)
vn('Theo Nguyên tắc dịch v2 — Nguyên tắc 2: References giữ nguyên tiếng Anh trong PDF gốc.', bold=True)
en('Bảng 1 (Study characteristics, p8-12) chiếm 5 trang trong PDF gốc với chi tiết 15 NC. Bảng 2 (Characteristics, training, supervision and role of PSWs, p15) tóm tắt cấu trúc PSW trong các NC. Cả 2 bảng giữ tiếng Anh trong PDF gốc.')
vn('Bảng 1 (Đặc điểm các nghiên cứu, trang 8-12) chiếm 5 trang trong PDF gốc với chi tiết 15 NC. Bảng 2 (Đặc điểm, đào tạo, giám sát và vai trò PSWs, trang 15) tóm tắt cấu trúc PSW. Cả 2 bảng GIỮ NGUYÊN TIẾNG ANH trong PDF gốc — không lặp lại trong bản dịch.')
d.add_paragraph()
en('References (p23-27) follow APA format and are retained in English in original PDF. See PDF for full reference list.')
vn('Phần References (trang 23-27) theo format APA và được giữ tiếng Anh trong PDF gốc. Xem PDF gốc để có danh mục tham khảo đầy đủ.')

d.add_paragraph()
foot = d.add_paragraph()
fr = foot.add_run('Bản dịch song ngữ Anh-Việt — biên soạn 29/04/2026. Số liệu (15 NC, 13 interventions, 8/13 nhắm distress nhẹ-vừa, 9/13 cho 16-25 tuổi) verified từ p1, p3, p7 PDF gốc. References + Tables giữ tiếng Anh (Nguyên tắc 2).')
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
