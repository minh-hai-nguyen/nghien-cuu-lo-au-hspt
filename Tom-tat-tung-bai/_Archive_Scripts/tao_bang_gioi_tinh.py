# -*- coding: utf-8 -*-
"""Bảng tổng hợp khác biệt giới tính liên bài — 11 bài NC"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GREEN = RGBColor(0, 0x80, 0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'; style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4); style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(2); s.right_margin = Cm(2)

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

# TIÊU ĐỀ
h = doc.add_heading('BẢNG TỔNG HỢP KHÁC BIỆT GIỚI TÍNH VỀ LO ÂU', 1)
for r in h.runs: r.font.name = 'Times New Roman'
h2 = doc.add_heading('So sánh 11 bài nghiên cứu — Phục vụ phần Thảo luận Đề cương', 2)
for r in h2.runs: r.font.name = 'Times New Roman'

doc.add_paragraph()

# BẢNG CHÍNH
headers = ['#', 'Tác giả, năm', 'Quốc gia', 'n', 'Công cụ', 'Nam (%)', 'Nữ (%)', 'P', 'Hướng', 'Bối cảnh đặc thù']
widths = [0.8, 3.0, 2.5, 1.8, 2.0, 1.5, 1.5, 1.5, 1.8, 4.0]

data = [
    ['01', 'Jenkins 2023', 'Hoa Kỳ (San Diego)', '75', 'PHQ-9A\nGAD-10', '—', '—', '0,002*\n0,016*', 'Nữ > Nam', 'Đa sắc tộc, THCS\n*PHQ-9A & GAD-10'],
    ['02', 'Saikia 2023', 'Ấn Độ\n(ĐB, Assam)', '360', 'DASS-21', '30,0%', '18,9%', '0,049', '⚠️ NAM > Nữ', 'Văn hóa bộ lạc\nĐông Bắc Ấn Độ'],
    ['03', 'Mandaknalli\n2021', 'Ấn Độ\n(đô thị tự trị)', '450', 'GAD-7', '41,7%*', '58,3%*', '0,022', 'Nữ > Nam', '*Trong nhóm lo âu\n(n=108)'],
    ['04', 'NSCH 2020', 'Hoa Kỳ\n(quốc gia)', '55.162', 'Chẩn đoán\nlâm sàng', '12,3%', '20,1%', '—', 'Nữ > Nam', 'Báo cáo cha mẹ\nchẩn đoán bác sĩ'],
    ['05', 'Alharbi 2019', 'Ả Rập Saudi\n(Qassim)', '1.245', 'GAD-7', '28,7%*', '71,3%*', '<0,001', 'Nữ > Nam', '*% trong nhóm\nlo âu nặng'],
    ['06', 'Nakie 2022', 'Ethiopia\n(Tây Bắc)', '810', 'DASS-21', '—', 'AOR\n1,304', '0,046', 'Nữ > Nam', 'AOR nữ vs nam\ncho trầm cảm'],
    ['07', 'Chen 2023', 'Trung Quốc\n(Tây, Tự Cống)', '63.205', 'PHQ-9\nGAD-7', '42%*', '58%*', '<0,001', 'Nữ > Nam', '*% trong nhóm\ncăng thẳng TT\nOR nữ = 1,55'],
    ['08', 'Wen 2020', 'Trung Quốc\n(nông thôn\nGiang Tây)', '900', 'MHT', '36,8%*', '63,2%*', '<0,001', 'Nữ > Nam', '*% nhóm lo âu nặng\nOR nam = 0,262\nNữ gấp ~4 lần'],
    ['09', 'Qiu 2022', 'Trung Quốc\n(Hợp Phì)', '2.079', 'SAS', '11,1%', '17,5%', '<0,01', 'Nữ > Nam', 'CES-D & SAS'],
    ['10', 'Xu 2021', 'Trung Quốc\n(Hà Nam)', '373.216', 'GAD-7', '10,11%', '9,66%', '<0,001', '⚠️ NAM > Nữ', 'COVID-19 đỉnh dịch\n02/2020, nông thôn\nOR nữ = 0,92 (BV)'],
    ['11', 'Bhardwaj 2020', 'Ấn Độ\n(Chandigarh)', '288', 'DASS-21', '36,0%*', '63,7%*', '<0,001', 'Nữ > Nam', '*Nặng+cực nặng\nChênh lệch lớn nhất'],
]

t = doc.add_table(rows=1+len(data), cols=len(headers))
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER

# Headers
for i, h_text in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h_text
    set_w(c, widths[i])
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(9)
    shade(c, 'D9E2F3')

# Data
for ri, rd in enumerate(data):
    for ci, v in enumerate(rd):
        c = t.rows[ri+1].cells[ci]
        set_w(c, widths[ci])
        # Clear default paragraph
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(v)
        r.font.name = 'Times New Roman'; r.font.size = Pt(8)

        # Color coding for direction column
        if ci == 8:  # Hướng
            if 'NAM > Nữ' in v:
                r.font.color.rgb = RED; r.bold = True
            elif 'Nữ > Nam' in v:
                r.font.color.rgb = BLUE

# Highlight rows with male > female
for row_idx in [2, 10]:  # Saikia (row 2), Xu (row 10) - 0-indexed data, +1 for header
    for ci in range(len(headers)):
        shade(t.rows[row_idx].cells[ci], 'FFF2CC')  # light yellow

doc.add_paragraph()

# PHÂN TÍCH TỔNG HỢP
h3 = doc.add_heading('PHÂN TÍCH TỔNG HỢP', 2)
for r in h3.runs: r.font.name = 'Times New Roman'

# Bảng thống kê
t2 = doc.add_table(rows=5, cols=3)
t2.style = 'Table Grid'; t2.alignment = WD_TABLE_ALIGNMENT.CENTER
stats_h = ['Xu hướng giới tính', 'Số bài', 'Bài cụ thể']
for i, h_text in enumerate(stats_h):
    c = t2.rows[0].cells[i]; c.text = h_text
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    shade(c, 'D9E2F3')

stats_data = [
    ['Nữ > Nam (phù hợp y văn)', '8/11 (72,7%)', '01, 03, 04, 05, 06, 07, 08, 09, 11'],
    ['NAM > Nữ (trái y văn)', '2/11 (18,2%)', '02 Saikia, 10 Xu'],
    ['Không so sánh / không rõ', '1/11 (9,1%)', '(không có)'],
    ['Tỷ lệ quốc tế (tham khảo)', '—', 'Nữ > Nam trong >90% NC\n(McLean et al., 2011)'],
]
for ri, rd in enumerate(stats_data):
    for ci, v in enumerate(rd):
        c = t2.rows[ri+1].cells[ci]; c.text = v
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if 'NAM > Nữ' in v:
                    r.font.color.rgb = RED; r.bold = True

doc.add_paragraph()

# THẢO LUẬN
h4 = doc.add_heading('THẢO LUẬN — GỢI Ý CHO ĐỀ CƯƠNG', 2)
for r in h4.runs: r.font.name = 'Times New Roman'

points = [
    ('1. Xu hướng chung phù hợp y văn quốc tế',
     '8/11 bài (72,7%) cho thấy nữ > nam về lo âu — phù hợp với tổng quan hệ thống của McLean et al. (2011) và Salk et al. (2017). Chênh lệch giới lớn nhất ở Bhardwaj 2020 (nữ 63,7% nặng+cực nặng vs nam 36,0%, P < 0,001) và Wen 2020 (OR nam = 0,262, nữ gấp gần 4 lần).'),
    ('2. Hai ngoại lệ quan trọng — nam > nữ',
     '• Saikia 2023 (Ấn Độ, n=360): nam 30,0% > nữ 18,9% (P = 0,049) — vùng Đông Bắc có văn hóa bộ lạc đặc thù. Tuy nhiên, P sát ngưỡng và các NC cùng khu vực (Kumar 2017, Nag 2019) cho thấy nữ > nam.\n• Xu 2021 (Trung Quốc, n=373.216): nam 10,11% > nữ 9,66% — chênh lệch rất nhỏ (0,45 điểm %), thu thập trong đỉnh dịch COVID-19 (02/2020). Bối cảnh khủng hoảng cấp tính có thể ảnh hưởng khác biệt giới.'),
    ('3. Yếu tố giải thích ngoại lệ',
     '• Văn hóa bộ lạc: kỳ vọng nam giới khác biệt ở Đông Bắc Ấn Độ (Saikia).\n• COVID-19: bối cảnh khủng hoảng cấp tính, nông thôn thiếu thông tin → nam lo lắng hơn (Xu).\n• Phương pháp: DASS-21 vs GAD-7 có thể đo lường khía cạnh khác nhau của lo âu.\n• Cỡ mẫu: Saikia n=360 (nhỏ, P=0,049 sát ngưỡng) vs Xu n=373.216 (rất lớn nhưng chênh lệch nhỏ).'),
    ('4. Hàm ý cho Đề tài tại Việt Nam',
     '• Cần kiểm tra xem bối cảnh Việt Nam (đặc biệt vùng dân tộc thiểu số) có ngoại lệ tương tự.\n• V-NAMHS 2022 không ghi nhận khác biệt giới rõ ràng — cần NC sâu hơn.\n• Hoa 2024 tại Hà Nội: nữ TB 1,74 > nam 1,50 (P < 0,01) — phù hợp xu hướng quốc tế.\n• Ngo Anh Vinh 2024 tại Lạng Sơn (dân tộc thiểu số): chưa phân tích giới chi tiết → GAP.'),
    ('5. Khuyến nghị nghiên cứu tiếp theo',
     '• NC riêng về khác biệt giới ở thanh thiếu niên DTTS vùng cao Việt Nam.\n• So sánh THCS vs THPT về xu hướng giới (Chen 2023 cho thấy yếu tố khác nhau).\n• Phân tích theo loại lo âu: nữ có thể cao hơn về lo âu xã hội, nam cao hơn về lo âu thi cử.'),
]

for title, body in points:
    p = doc.add_paragraph()
    r = p.add_run(title)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.font.color.rgb = BLUE

    p2 = doc.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)

doc.save('BANG_GIOI_TINH_LIEN_BAI.docx')
sys.stderr.write('BANG_GIOI_TINH_LIEN_BAI.docx OK\n')
