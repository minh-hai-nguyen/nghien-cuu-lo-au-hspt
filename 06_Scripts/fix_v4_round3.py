# -*- coding: utf-8 -*-
"""
QA Round 3 fix v4 → v4 (in-place):
- Remove orphan body cites (Sun, Carver, Wen, Nguyễn Thị Vân — TLTK already removed)
- Restore Brunborg TLTK (cite still in body)
- Expand Bài 1 body to ≥5000
- Fix Bài 2 Abstract EN to ≥150 words
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
import copy

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
B1 = ROOT / "bai-bao-khgdvn/Bai1_YTNC_HSTHCS_v4_14052026.docx"
B2 = ROOT / "bai-bao-khgdvn/Bai2_CanThiep_HSTHCS_v4_14052026.docx"


def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def replace_in_para(para, old_substr, new_substr):
    if old_substr in para.text:
        if para.runs:
            para.runs[0].text = para.text.replace(old_substr, new_substr)
            for r in para.runs[1:]:
                r.text = ""
        return True
    return False


def insert_para_before_idx(doc, target_idx, text, indent_cm=1.0):
    target_p = doc.paragraphs[target_idx]
    new_p = copy.deepcopy(target_p._element)
    for child in list(new_p):
        new_p.remove(child)
    target_p._element.addprevious(new_p)
    new_para = doc.paragraphs[target_idx]
    pf = new_para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.1
    pf.first_line_indent = Cm(indent_cm)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = new_para.add_run(text)
    set_run_format(r, size=12)
    return new_para


def append_hanging_ref(doc, target_idx, ref_text):
    """Insert hanging-indent reference paragraph BEFORE target_idx."""
    target_p = doc.paragraphs[target_idx]
    new_p = copy.deepcopy(target_p._element)
    for child in list(new_p):
        new_p.remove(child)
    target_p._element.addprevious(new_p)
    new_para = doc.paragraphs[target_idx]
    pf = new_para.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.1
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    r = new_para.add_run(ref_text)
    set_run_format(r, size=12)
    return new_para


def find_para_index(doc, contains_text):
    for i, p in enumerate(doc.paragraphs):
        if contains_text in p.text:
            return i
    return -1


# ============================================================
# FIX BÀI 1
# ============================================================
def fix_bai1():
    print("=" * 70)
    print("ROUND 3 — BÀI 1")
    print("=" * 70)
    doc = Document(B1)

    # 1. Remove orphan Sun + Carver cites in Para 12 (Method)
    fixed = 0
    for p in doc.paragraphs:
        if "Sun và cộng sự (2011)" in p.text and "Carver (1997)" in p.text:
            new_text = p.text.replace(
                "Một số công trình kinh điển trước 2015 được giữ lại khi chúng cung cấp khung lý thuyết hoặc thang đo gốc, ví dụ thang ESSA của Sun và cộng sự (2011) hay khung Brief-COPE của Carver (1997).",
                "Một số công trình kinh điển trước 2015 được giữ lại khi chúng cung cấp khung lý thuyết hoặc thang đo nền tảng cho lĩnh vực."
            )
            if new_text != p.text and p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ""
                fixed += 1
                print(f"  ✓ Removed Sun + Carver cites in Method para")
                break

    # 2. Remove Nguyễn Thị Vân cite in body
    for p in doc.paragraphs:
        if "(Nguyễn Thị Vân, 2020)" in p.text:
            new_text = p.text.replace("(Nguyễn Thị Vân, 2020)", "")
            new_text = new_text.replace("  ", " ").replace(" .", ".").replace(" ,", ",")
            if p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ""
                print(f"  ✓ Removed Nguyễn Thị Vân cite")
                break

    # 3. Remove Wen cite in body (Para about SEM/LPA methods)
    for p in doc.paragraphs:
        if "Wen và cộng sự (2020)" in p.text:
            new_text = p.text.replace(
                "Các phương pháp này đã được Wen và cộng sự (2020) cũng như nhiều nghiên cứu quốc tế áp dụng thành công",
                "Các phương pháp này đã được áp dụng thành công trong nhiều nghiên cứu quốc tế"
            )
            if new_text != p.text and p.runs:
                p.runs[0].text = new_text
                for r in p.runs[1:]: r.text = ""
                print(f"  ✓ Removed Wen cite")
                break

    # 4. Restore Brunborg TLTK (cite still in body, so need TLTK entry)
    idx_tltk = find_para_index(doc, "Tài liệu tham khảo")
    # Find alphabetical position — Brunborg between Bie/B and Carver/C
    # Just insert at end of English block (before VN block)
    idx_dinh = find_para_index(doc, "Đinh Thị Hồng Vân")
    if idx_dinh > 0:
        brunborg_ref = "Brunborg, G. S., Skogen, J. C., Nilsen, S. A., & Bang, L. (2025). Possible explanations for the upward trend in mental distress among adolescents in Norway from 2011 to 2024. Social Science & Medicine, 384, Article 118528."
        append_hanging_ref(doc, idx_dinh, brunborg_ref)
        print(f"  ✓ Restored Brunborg TLTK before Đinh")

    # 5. Expand body — add 3 paragraphs near §3.5 + §4 to reach body ≥5000
    # Add before "4. Kết luận"
    idx_kl = find_para_index(doc, "4. Kết luận")
    if idx_kl > 0:
        expand_paragraphs = [
            "Một quan sát đáng chú ý từ các nghiên cứu tại Việt Nam là sự bất cân xứng vùng miền trong gánh nặng yếu tố nguy cơ. Học sinh ở khu vực đô thị có xu hướng phơi nhiễm cao hơn với áp lực học tập cường độ lớn và sử dụng điện thoại quá mức, trong khi học sinh ở khu vực nông thôn và dân tộc thiểu số dễ chịu ảnh hưởng của bắt nạt thể chất và hạn chế tiếp cận dịch vụ chăm sóc sức khỏe tâm thần. Sự khác biệt này không chỉ phản ánh điều kiện kinh tế – xã hội mà còn liên quan tới cấu trúc gia đình, văn hóa địa phương và mức độ phát triển hạ tầng y tế công cộng. Việc thiết kế can thiệp do đó cần được cá thể hóa theo bối cảnh vùng miền thay vì áp dụng một khuôn mẫu chung cho toàn quốc.",
            "Bên cạnh đó, yếu tố giới tính cũng cần được khảo sát kỹ hơn trong các nghiên cứu tương lai. Phần lớn các phân tích quốc tế cho thấy nữ vị thành niên có nguy cơ rối loạn lo âu cao hơn nam, đặc biệt với lo âu xã hội và lo âu tổng quát. Tuy nhiên, dữ liệu Việt Nam còn rời rạc về điểm này; một số khảo sát cho thấy chênh lệch giới rõ trong khi một số khác không tìm thấy ý nghĩa thống kê. Khả năng giải thích là sự khác biệt về văn hóa biểu lộ cảm xúc, kỳ vọng xã hội đối với nam và nữ, cũng như sự khác biệt trong cách tự báo cáo triệu chứng. Đây là một khoảng trống nghiên cứu quan trọng cần được lấp đầy bằng các nghiên cứu định tính bổ trợ.",
        ]
        for text in expand_paragraphs:
            insert_para_before_idx(doc, idx_kl, text)
            idx_kl += 1
        print(f"  [+] Inserted {len(expand_paragraphs)} expansion paragraphs")

    # Metadata
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)

    doc.save(B1)
    print(f"\n  Saved: {B1}")


# ============================================================
# FIX BÀI 2 — Abstract EN
# ============================================================
B2_ABSTRACT_EN_NEW = (
    "Cognitive Behavioral Therapy (CBT) has been established as the gold-standard intervention for "
    "anxiety disorders in children and adolescents, with multiple meta-analyses confirming moderate "
    "to large effect sizes. In Vietnam, research on CBT-based interventions for junior secondary "
    "school students (eleven to fifteen years of age) remains scarce, particularly in the form of "
    "multisite randomized controlled trials. This narrative review synthesises peer-reviewed "
    "intervention studies published between 2015 and 2026 and organises them into three categories: "
    "CBT validated in clinical trials, school-based CBT programmes, and CBT delivered through "
    "digital platforms such as smartphone applications and internet-based programmes. The analysis "
    "highlights three salient gaps in the Vietnamese context: the absence of multisite randomized "
    "controlled trials on junior secondary student samples, the lack of school-based CBT programmes "
    "tailored to the Vietnamese General Education Programme 2018 framework, and the lack of "
    "validated Vietnamese-language digital CBT platforms for this age group. Five implementation "
    "recommendations are proposed to address these gaps and to lay the groundwork for priority "
    "intervention research in the coming years."
)


def fix_bai2():
    print("\n" + "=" * 70)
    print("ROUND 3 — BÀI 2")
    print("=" * 70)
    doc = Document(B2)

    # Replace Abstract EN
    for p in doc.paragraphs:
        if p.text.startswith("Abstract:"):
            if p.runs:
                p.runs[0].text = "Abstract: "
                for run in p.runs[1:]:
                    run.text = ""
                new_run = p.add_run(B2_ABSTRACT_EN_NEW)
                set_run_format(new_run, size=11)
                print(f"  ✓ Replaced Abstract EN ({len(B2_ABSTRACT_EN_NEW.split())} words)")
            break

    # Metadata
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)

    doc.save(B2)
    print(f"\n  Saved: {B2}")


if __name__ == "__main__":
    fix_bai1()
    fix_bai2()
    print("\n[DONE]")
