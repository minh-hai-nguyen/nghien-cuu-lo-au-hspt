# -*- coding: utf-8 -*-
"""Batch doi thuat ngu 'lan toa' -> 'tong quat' cho GAD theo confirm tu thay.
Update 15 files + backup + clean metadata.
01/06/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Cac file can update
FILES = [
    'Luận án TS/01_LuanAn_v4_ChuanTrinhBay_28052026.docx',
    'Luận án TS/TomTatLA_v2_VERIFIED_29052026.docx',
    'Luận án TS/TomTatLA_EN_v1_29052026.docx',
    'Luận án TS/TrichYeuLA_CongThiHang_v2_29052026.docx',
    'Luận án TS/DanhMucCongTrinh_EN_v1_29052026.docx',
    'bai-bao-Q1/BoCuc_Q1_Q3_PhuongAnA_v2_01062026.docx',
    'bai-bao-Q1/RaSoat_Q1_01062026.docx',
    'bai-bao-Q1/RaSoat_Q3_01062026.docx',
    'bai-bao-Q1/IssuesPriority_Q1Q3_01062026.docx',
    'bai-bao-Q1/Outline_Q1_v3_01062026.docx',
    'bai-bao-Q1/Outline_Q3_v3_01062026.docx',
    'bai-bao-Q1/OutlineBilingual_Q1_01062026.docx',
    'bai-bao-Q1/OutlineBilingual_Q3_01062026.docx',
    'bai-bao-Q1/BaiA_JAD_OUTLINE_v1_30052026.docx',
    'bai-bao-Q1/BaiD_StressHealth_OUTLINE_v1_30052026.docx',
]

# Cac thay the (theo thu tu de phrase dai duoc xu ly truoc phrase ngan)
REPLACEMENTS = [
    # Specific phrases first (longer first)
    ('rối loạn lo âu lan toả', 'rối loạn lo âu tổng quát'),
    ('rối loạn lo âu lan tỏa', 'rối loạn lo âu tổng quát'),
    ('Rối loạn lo âu lan toả', 'Rối loạn lo âu tổng quát'),
    ('Rối loạn lo âu lan tỏa', 'Rối loạn lo âu tổng quát'),
    ('RỐI LOẠN LO ÂU LAN TỎA', 'RỐI LOẠN LO ÂU TỔNG QUÁT'),
    ('RỐI LOẠN LO ÂU LAN TOẢ', 'RỐI LOẠN LO ÂU TỔNG QUÁT'),
    # "lo âu lan tỏa" stand-alone
    ('lo âu lan toả', 'lo âu tổng quát'),
    ('lo âu lan tỏa', 'lo âu tổng quát'),
    ('Lo âu lan toả', 'Lo âu tổng quát'),
    ('Lo âu lan tỏa', 'Lo âu tổng quát'),
    ('LO ÂU LAN TỎA', 'LO ÂU TỔNG QUÁT'),
    ('LO ÂU LAN TOẢ', 'LO ÂU TỔNG QUÁT'),
    # Viết tắt
    ('RLLALT', 'RLLATQ'),
    # Cụm khác
    ('(RLLALT', '(RLLATQ'),
    ('RLLALT)', 'RLLATQ)'),
    ('rối loạn lo âu lan toa', 'rối loạn lo âu tổng quát'),  # without diacritic
]


def replace_in_runs(paragraph):
    """Replace text in runs preserving formatting."""
    count = 0
    for run in paragraph.runs:
        for old, new in REPLACEMENTS:
            if old in run.text:
                run.text = run.text.replace(old, new)
                count += 1
    return count


def replace_full_text(paragraph):
    """If run-level replace doesn't work (text split across runs), try full text."""
    full = paragraph.text
    has_match = any(old in full for old, _ in REPLACEMENTS)
    if not has_match:
        return 0
    # Replace whole paragraph text
    new_full = full
    for old, new in REPLACEMENTS:
        new_full = new_full.replace(old, new)
    if new_full == full:
        return 0

    # Apply: clear runs + add one new run with replaced text (loses format)
    # OR: smart approach - check if first run captures full text or each run is separate
    if len(paragraph.runs) == 1:
        paragraph.runs[0].text = new_full
        return 1

    # Multi-run case: try replacing in each run first
    # If text spans across runs (replacement string broken), use brute approach
    # For now, just do run-level replacement (handled by replace_in_runs)
    # Return 0 to indicate not fully handled here
    return 0


def process_file(rel_path):
    """Process one file: backup, replace, save."""
    full_path = os.path.join(ROOT, rel_path)
    if not os.path.exists(full_path):
        return {'file': rel_path, 'status': 'NOT_FOUND', 'replacements': 0}

    # Backup
    backup = full_path.replace('.docx', '_BEFORE_GAD_DOI_01062026.docx')
    shutil.copy(full_path, backup)

    try:
        d = Document(full_path)
    except Exception as e:
        return {'file': rel_path, 'status': f'READ_ERR: {e}', 'replacements': 0}

    total = 0

    # Paragraphs
    for para in d.paragraphs:
        total += replace_in_runs(para)
        # Also try full-text replace for cases where run-level missed
        full = para.text
        has = any(old in full for old, _ in REPLACEMENTS)
        if has:
            # Check if already replaced via runs
            new_full = full
            for old, new in REPLACEMENTS:
                new_full = new_full.replace(old, new)
            # If full text after runs still has old, do paragraph-level (loses format)
            if any(old in para.text for old, _ in REPLACEMENTS):
                # Multi-run case - text split. Apply paragraph-level rebuild
                if para.runs:
                    # Combine all into first run, replace
                    combined = para.text
                    for old, new in REPLACEMENTS:
                        combined = combined.replace(old, new)
                    # Get format from first run
                    first = para.runs[0]
                    fn = first.font.name
                    fs = first.font.size
                    fb = first.bold
                    fi = first.italic
                    # Clear all runs
                    for r in list(para.runs):
                        r._element.getparent().remove(r._element)
                    new_run = para.add_run(combined)
                    new_run.font.name = fn
                    if fs:
                        new_run.font.size = fs
                    new_run.bold = fb
                    new_run.italic = fi
                    total += 1

    # Tables
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    total += replace_in_runs(para)
                    # Same multi-run handling
                    if any(old in para.text for old, _ in REPLACEMENTS):
                        if para.runs:
                            combined = para.text
                            for old, new in REPLACEMENTS:
                                combined = combined.replace(old, new)
                            first = para.runs[0]
                            fn = first.font.name
                            fs = first.font.size
                            fb = first.bold
                            fi = first.italic
                            for r in list(para.runs):
                                r._element.getparent().remove(r._element)
                            new_run = para.add_run(combined)
                            new_run.font.name = fn
                            if fs:
                                new_run.font.size = fs
                            new_run.bold = fb
                            new_run.italic = fi
                            total += 1

    # Clean metadata
    cp = d.core_properties
    cp.author = ''; cp.title = ''; cp.subject = ''
    cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
    cp.category = ''; cp.identifier = ''; cp.version = ''

    d.save(full_path)
    return {'file': rel_path, 'status': 'OK', 'replacements': total,
            'backup': os.path.basename(backup)}


# Process all files
print('='*70)
print('BATCH UPDATE: "lan tỏa" → "tổng quát" cho GAD')
print('='*70)
print()

results = []
for f in FILES:
    r = process_file(f)
    results.append(r)
    icon = '✓' if r['status'] == 'OK' else '✗' if 'ERR' in r['status'] else '?'
    print(f'  {icon} {r["replacements"]:>4} replacements | {f}')
    if 'backup' in r:
        print(f'         backup: {r["backup"]}')

# Summary
print()
print('='*70)
print(f'TỔNG: {sum(r["replacements"] for r in results)} replacements across '
      f'{sum(1 for r in results if r["status"] == "OK")} files')
print('='*70)

# Verify
print()
print('=== VERIFY no "lan tỏa" remaining in GAD context ===')
for f in FILES:
    full = os.path.join(ROOT, f)
    if not os.path.exists(full):
        continue
    d = Document(full)
    text = '\n'.join(p.text for p in d.paragraphs)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                text += '\n' + c.text
    # Check
    has_old = ('lo âu lan tỏa' in text.lower() or 'lo âu lan toả' in text.lower()
               or 'RLLALT' in text)
    has_new = ('lo âu tổng quát' in text.lower() or 'RLLATQ' in text)
    icon = '✓' if has_new and not has_old else ('⚠' if has_old else '—')
    print(f'  {icon} new={has_new} old={has_old} | {f}')
