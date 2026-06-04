# -*- coding: utf-8 -*-
"""QT21-24: Norway, ScreenTime, JAACAP, WHO Europe — chuẩn CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc
def rb(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(doc, t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== QT21: NORWAY 2011-2024 =====
d = make_doc()
bl(d, 'Tóm tắt bài QT-21', bold=True)
rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d, 'Công trình nghiên cứu \u00ab Giải thích khả thi cho xu hướng tăng căng thẳng tâm thần ở thanh thiếu niên Na Uy từ 2011 đến 2024 \u00bb (Possible Explanations for the Upward Trend in Mental Distress among Adolescents in Norway from 2011 to 2024), xuất bản trên Social Science & Medicine, 2025 (Q1, IF \u2248 5,4). 10 trang. Phân tích xu hướng 13 năm dữ liệu lặp lại từ khảo sát quốc gia Na Uy.')

rb(d, 'Phương pháp nghiên cứu')
bl(d, 'Công trình sử dụng phân tích decomposition (phân rã đóng góp) trên dữ liệu cắt ngang lặp lại 2011\u20132024 từ khảo sát quốc gia VTN Na Uy. Nói cách khác, phương pháp này cho phép tách biệt đóng góp của từng yếu tố (bất mãn trường học, mạng xã hội, v.v.) vào xu hướng tăng tổng thể.', bold=True)

bl(d, 'Tổng quan tài liệu của chúng tôi cho thấy xu hướng tăng căng thẳng tâm thần ở VTN đã được báo cáo tại nhiều nước phương Tây (JAACAP 2024: lo âu tăng gấp đôi tại Mỹ; NSCH 2020: tăng 61%), nhưng các yếu tố giải thích xu hướng này vẫn chưa được làm rõ. Nghiên cứu này ở Na Uy sử dụng phương pháp decomposition để giải đáp câu hỏi đó.', bold=True)

rb(d, 'Kết quả nghiên cứu định lượng')
tbl(d, ['Yếu tố giải thích', 'Đóng góp vào xu hướng tăng', 'Ý nghĩa'],
    [['Bất mãn với trường học', 'Lớn nhất', 'Áp lực học tập, môi trường trường học'],
     ['Thời gian dùng mạng xã hội', 'Đáng kể', 'Phù hợp Nature 2025, Chen 2023'],
     ['Giới tính', 'Cả nam và nữ đều tăng', 'Nữ tăng nhanh hơn'],
     ['Xu hướng tổng thể 2011\u20132024', 'Tăng liên tục 13 năm', '\u2014']],
    widths=[4.5, 3.5, 4.0])
d.add_paragraph()

bh3(d, 'Đối chiếu liên bài : Xu hướng tăng lo âu toàn cầu được xác nhận bởi nhiều nguồn: JAACAP 2024 (Mỹ: 9,6% \u2192 19,2%, 8 năm), NSCH 2020 (Mỹ: tăng 61%, 7 năm), Ireland 2024 (tăng 2012\u20132019), Korea 2024 (tăng sau COVID). Nghiên cứu Na Uy bổ sung bằng xác định NGUYÊN NHÂN: bất mãn trường học và mạng xã hội.')

rb(d, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d, '*Bất mãn trường học.* Yếu tố giải thích mạnh nhất cho xu hướng tăng, phù hợp với Wen 2020 (áp lực học tập OR = 11,58) và Anderson 2025 (48/52 NC liên kết áp lực học tập với SKTT kém).')
bl(d, '*Mạng xã hội.* Đóng góp đáng kể, phù hợp Nature Human Behaviour 2025 (VTN có rối loạn SKTT dùng MXH nhiều hơn) và JAMA 2024 (screen time ảnh hưởng SKTT trong RCT).')

rb(d, 'Kết luận')
bl(d, 'Dữ liệu của nghiên cứu này, cho thấy xu hướng tăng liên tục 13 năm ở căng thẳng tâm thần VTN Na Uy, với bất mãn trường học và mạng xã hội là hai yếu tố giải thích chính, gợi ý rằng can thiệp cần nhắm vào cải thiện môi trường trường học và quản lý thời gian mạng xã hội — không chỉ tại Na Uy mà có thể áp dụng cho bối cảnh toàn cầu bao gồm Việt Nam.', bold=True)

rh2(d, 'Phản biện')
bl(d, 'Social Science & Medicine Q1 IF = 5,4, xu hướng 13 năm, decomposition nâng cao. Tuy nhiên, chỉ Na Uy \u2014 khác biệt văn hóa lớn với châu Á. Đo "mental distress" chung, không tách lo âu riêng. Không phải nghiên cứu dọc theo cùng cá nhân.')

rh2(d, 'Hướng nghiên cứu tiếp theo')
bl(d, 'Áp dụng phân tích decomposition cho dữ liệu VN (Hoàng Trung Học 2025 có 2 thời điểm). So sánh xu hướng Na Uy vs VN. Can thiệp giảm bất mãn trường học + kiểm soát MXH.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50 Cao. Q1 IF = 5,4, xu hướng 13 năm, decomposition xác định nguyên nhân.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.save('QT21_Norway_2025.docx')
sys.stderr.write('QT21 OK\n')

# ===== QT22: SCREEN TIME BJCP =====
d = make_doc()
bl(d, 'Tóm tắt bài QT-22', bold=True)
rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d, 'Công trình nghiên cứu \u00ab Mối liên quan cắt ngang và dọc giữa thời gian sử dụng màn hình với trầm cảm và lo âu ở thanh thiếu niên \u00bb (Cross-sectional and Longitudinal Associations of Screen Time with Adolescent Depression and Anxiety), do Li SH và cộng sự (2025), xuất bản trên British Journal of Clinical Psychology, 64:873\u2013887 (Q1, IF \u2248 3,0). 15 trang.')

rb(d, 'Phương pháp nghiên cứu')
bl(d, 'Công trình sử dụng thiết kế DỌC (longitudinal) — theo dõi cùng nhóm VTN theo thời gian — kết hợp phân tích cắt ngang. Nói cách khác, đây là một trong số ít nghiên cứu có thể suy luận chiều hướng nhân quả: thời gian màn hình DỰ BÁO trầm cảm và lo âu sau 1 năm, không chỉ tương quan cùng thời điểm.', bold=True)

bl(d, 'Tổng quan tài liệu cho thấy đa số NC về screen time và SKTT là cắt ngang (không suy luận nhân quả). Nghiên cứu này lấp đầy khoảng trống đó.', bold=True)

rb(d, 'Kết quả nghiên cứu định lượng')
tbl(d, ['Thiết kế', 'Phát hiện chính', 'So sánh liên bài'],
    [['Cắt ngang', 'Screen time liên quan trầm cảm + lo âu', 'Phù hợp Chen 2023 (game OR = 5,00)'],
     ['DỌC', 'Screen time DỰ BÁO trầm cảm/lo âu sau 1 năm', 'Vượt trội: bằng chứng nhân quả'],
     ['Cơ chế', 'Tăng screen time \u2192 tăng triệu chứng SKTT', 'Hoàng Trung Học VN: Beta = 0,176'],
     ['Ý nghĩa', 'Giảm screen time có thể cải thiện SKTT', 'JAMA 2024: RCT xác nhận']],
    widths=[3.0, 5.0, 4.5])
d.add_paragraph()

bh3(d, 'Đối chiếu liên bài : Thiết kế dọc của nghiên cứu này bổ sung bằng chứng nhân quả cho mối quan hệ screen time \u2192 SKTT mà Chen 2023 (OR = 5,00 cho game), Hoàng Trung Học 2025 (Beta = 0,176 cho điện tử) và Norway 2025 (MXH giải thích xu hướng tăng) đã gợi ý từ thiết kế cắt ngang.')

rb(d, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d, '*Bằng chứng nhân quả.* Đây là điểm mạnh vượt trội — thiết kế dọc cho phép kết luận screen time GÂY RA (không chỉ liên quan) tăng trầm cảm/lo âu.')
bl(d, '*Ứng dụng can thiệp.* Kết hợp với JAMA 2024 (RCT giảm screen time cải thiện SKTT), nghiên cứu này cung cấp cơ sở mạnh cho can thiệp quản lý thời gian màn hình tại trường học VN.')

rb(d, 'Kết luận')
bl(d, 'Dữ liệu dọc cho thấy thời gian sử dụng màn hình dự báo tăng triệu chứng trầm cảm và lo âu ở VTN sau 1 năm, gợi ý rằng can thiệp giảm screen time là chiến lược phòng ngừa có cơ sở khoa học mạnh.', bold=True)

rh2(d, 'Phản biện')
bl(d, 'BJCP Q1, thiết kế DỌC vượt trội. Tuy nhiên, cần đọc chi tiết cỡ mẫu, công cụ đo lo âu cụ thể. Chỉ 1 bối cảnh văn hóa — cần xác nhận ở châu Á.')

rh2(d, 'Hướng nghiên cứu tiếp theo')
bl(d, 'Nghiên cứu dọc tương tự tại VN. RCT giảm screen time tại trường THCS/THPT VN. So sánh bối cảnh phương Tây vs châu Á.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50 Cao. BJCP Q1, thiết kế dọc — bằng chứng nhân quả.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.save('QT22_ScreenTime_2025.docx')
sys.stderr.write('QT22 OK\n')

# ===== QT23: JAACAP US TRENDS =====
d = make_doc()
bl(d, 'Tóm tắt bài QT-23', bold=True)
rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d, 'Công trình nghiên cứu \u00ab Xu hướng rối loạn tâm thần ở trẻ em và thanh thiếu niên nhận điều trị trong hệ thống sức khỏe tâm thần công lập Hoa Kỳ \u00bb (Trends in Mental Disorders in Children and Adolescents Receiving Treatment in the State Mental Health System), xuất bản trên Journal of the American Academy of Child & Adolescent Psychiatry (JAACAP), 2024 (Q1, IF \u2248 11,0). Dữ liệu Mental Health Client-Level Data 2013\u20132021, tổng số 13.684.154 hồ sơ. 15 trang.')

rb(d, 'Phương pháp nghiên cứu')
bl(d, 'Công trình phân tích xu hướng tỷ lệ các rối loạn tâm thần được chẩn đoán trong hệ thống y tế công lập Mỹ, sử dụng dữ liệu hành chính quốc gia 8 năm. Nói cách khác, đây là phân tích xu hướng dựa trên chẩn đoán lâm sàng (không phải sàng lọc) từ hệ thống dịch vụ thực tế.', bold=True)

rb(d, 'Kết quả nghiên cứu định lượng')
bl(d, 'Xu hướng rối loạn tâm thần ở trẻ em/VTN Mỹ 2013\u20132021 (N = 13.684.154 hồ sơ):', bold=True)
tbl(d, ['Rối loạn', '2013', '2021', 'AOR (KTC 95%)', 'Xu hướng'],
    [['Rối loạn lo âu', '9,6%', '19,2%', '2,17 (1,85\u20132,55)***', '\u2191 TĂNG GẤP ĐÔI'],
     ['Rối loạn trầm cảm', '13,4%', '17,0%', '1,20 (1,03\u20131,41)*', '\u2191 Tăng'],
     ['Sang chấn/stress', '22,7%', '27,4%', '1,31 (1,09\u20131,57)**', '\u2191 Tăng'],
     ['Rối loạn lưỡng cực', '\u2014', '\u2014', '\u2014', '\u2193 Giảm'],
     ['Rối loạn hành vi/chống đối', '\u2014', '\u2014', '\u2014', '\u2193 Giảm']],
    widths=[3.5, 1.5, 1.5, 3.5, 2.5])
d.add_paragraph()

bh3(d, 'Đối chiếu liên bài : Lo âu TĂNG GẤP ĐÔI (9,6% \u2192 19,2%, AOR = 2,17) trong 8 năm — xu hướng mạnh nhất. So sánh: NSCH 2020 (lo âu tăng 61% tại Mỹ, 2016\u20132023), Norway 2025 (tăng liên tục 13 năm), Korea 2024 (tăng sau COVID). Xu hướng nhất quán toàn cầu.')

rb(d, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d, '*Lo âu — rối loạn tăng nhanh nhất.* AOR = 2,17 vượt trội so với trầm cảm (1,20) và sang chấn (1,31). Nói cách khác, lo âu đang trở thành cuộc khủng hoảng SKTT hàng đầu ở trẻ em/VTN phương Tây.')
bl(d, '*Rối loạn hành vi giảm.* Lưỡng cực và hành vi/chống đối giảm — có thể phản ánh thay đổi thực hành chẩn đoán (từ hành vi hóa sang nội hóa) chứ không chỉ thay đổi tỷ lệ thực.')
bl(d, '*Dữ liệu chẩn đoán lâm sàng.* Khác với sàng lọc (Hoa 2024: 40,6% GAD-7), dữ liệu từ hệ thống điều trị phản ánh tỷ lệ được chẩn đoán thực tế. So sánh với V-NAMHS 2022 (DISC-5: 2,3% chẩn đoán) cho thấy khoảng cách tiếp cận dịch vụ khổng lồ giữa Mỹ và VN.')

rb(d, 'Kết luận')
bl(d, 'Dữ liệu từ 13,7 triệu hồ sơ lâm sàng Mỹ, cho thấy lo âu tăng gấp đôi (AOR = 2,17) trong 8 năm — xu hướng mạnh nhất trong tất cả rối loạn tâm thần ở trẻ em, gợi ý rằng lo âu đang trở thành ưu tiên SKTT hàng đầu toàn cầu. Tại Việt Nam, nơi chỉ 8,4% VTN tiếp cận dịch vụ, tỷ lệ thực có thể cao hơn nhiều so với báo cáo.', bold=True)

rh2(d, 'Phản biện')
bl(d, 'JAACAP Q1 IF = 11 — tạp chí hàng đầu. n = 13,7 triệu hồ sơ lâm sàng, xu hướng 8 năm. Tuy nhiên, chỉ trẻ ĐANG điều trị — thiên lệch chọn (không bao gồm trẻ chưa chẩn đoán). Xu hướng tăng có thể phản ánh tăng phát hiện + tăng thực sự cùng lúc. Chỉ Mỹ — bối cảnh dịch vụ SKTT phát triển, khác VN.')

rh2(d, 'Hướng nghiên cứu tiếp theo')
bl(d, 'Phân tích xu hướng tương tự từ dữ liệu Bệnh viện Nhi TW VN. So sánh hệ thống y tế Mỹ (82,6% tiếp cận) vs VN (8,4%). Đánh giá nguyên nhân xu hướng tăng (phát hiện vs thực sự).')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. JAACAP Q1 IF = 11, n = 13,7 triệu, xu hướng 8 năm, dữ liệu chẩn đoán lâm sàng.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.save('QT23_JAACAP_US_2024.docx')
sys.stderr.write('QT23 OK\n')

# ===== QT24: WHO EUROPE LANCET =====
d = make_doc()
bl(d, 'Tóm tắt bài QT-24', bold=True)
rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d, 'Công trình \u00ab Sức khỏe tâm thần của trẻ em và thanh niên trong khu vực WHO châu Âu \u00bb (Mental Health of Children and Young People in the WHO Europe Region), xuất bản trên The Lancet Regional Health \u2014 Europe, 2025 (Q1, IF \u2248 15,0). 13 trang. Tổng quan toàn diện cho chính sách.')

rb(d, 'Phương pháp nghiên cứu')
bl(d, 'Tổng quan chính sách (policy review) tổng hợp dữ liệu từ nhiều nguồn trong khu vực WHO châu Âu. Nói cách khác, đây không phải nghiên cứu gốc mà là bài phân tích tổng hợp phục vụ hoạch định chính sách.', bold=True)

rb(d, 'Kết quả nghiên cứu định lượng')
tbl(d, ['Chỉ số', 'Dữ liệu', 'Nguồn'],
    [['VTN có rối loạn SKTT', '9 triệu (10\u201319 tuổi) châu Âu', 'WHO Europe'],
     ['Lo âu + trầm cảm', '>50% tất cả ca rối loạn', '\u2014'],
     ['Xu hướng', 'Tăng, đặc biệt sau COVID-19', 'Nhiều khảo sát quốc gia'],
     ['Giới tính', 'Nữ bị ảnh hưởng nhiều hơn', 'Phù hợp y văn'],
     ['Yếu tố nguy cơ', 'Mạng xã hội, áp lực học tập, nghèo đói', 'Liên quốc gia'],
     ['Khuyến nghị', 'Tích hợp SKTT vào giáo dục + y tế', 'Chính sách WHO']],
    widths=[4.0, 4.5, 3.0])
d.add_paragraph()

bh3(d, 'Đối chiếu liên bài : Con số 9 triệu VTN châu Âu có rối loạn SKTT cho thấy quy mô vấn đề rất lớn. So sánh: ASEAN có 80,4 triệu ca rối loạn tâm thần (GBD 2025) — gánh nặng lớn hơn gấp nhiều lần do dân số đông hơn và dịch vụ SKTT yếu hơn. Tại VN, chỉ 8,4% tiếp cận dịch vụ so với hệ thống y tế phát triển của châu Âu.')

rb(d, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d, '*Khuyến nghị tích hợp SKTT vào giáo dục.* Phù hợp với phát hiện của Wen 2020 (hỗ trợ SKTT trường OR = 0,562 bảo vệ) và đề xuất của Zhameden 2025 (cần can thiệp tại trường).')
bl(d, '*Lo âu + trầm cảm > 50% ca rối loạn.* Xác nhận lo âu là vấn đề hàng đầu, phù hợp JAACAP 2024 (lo âu tăng gấp đôi) và tất cả NC VN (lo âu luôn cao nhất trong 3 chỉ số DAS).')

rb(d, 'Kết luận')
bl(d, 'Dữ liệu tổng hợp WHO châu Âu, cho thấy 9 triệu VTN đang sống với rối loạn SKTT và lo âu chiếm > 50% tổng ca, gợi ý rằng tích hợp SKTT vào hệ thống giáo dục và y tế là chiến lược thiết yếu — bài học cho Việt Nam trong bối cảnh xây dựng chương trình SKTT học đường.', bold=True)

rh2(d, 'Phản biện')
bl(d, 'Lancet Regional Health Europe Q1 IF = 15 — tạp chí uy tín rất cao. Phạm vi toàn châu Âu. Tuy nhiên, tổng quan chính sách, không phải NC gốc — thiếu số liệu cụ thể từng nước. Bối cảnh châu Âu (dịch vụ SKTT phát triển) khác biệt lớn với ASEAN/VN.')

rh2(d, 'Hướng nghiên cứu tiếp theo')
bl(d, 'So sánh mô hình can thiệp châu Âu (school-based) với ASEAN/VN. Đánh giá khả năng áp dụng khuyến nghị WHO cho bối cảnh VN. Tổng quan tương tự cho khu vực ASEAN.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. Lancet Q1 IF = 15, phạm vi toàn châu Âu, hướng chính sách.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.save('QT24_WHO_Europe_2025.docx')
sys.stderr.write('QT24 OK\n')
