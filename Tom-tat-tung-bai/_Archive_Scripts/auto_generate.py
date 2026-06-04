# -*- coding: utf-8 -*-
"""
Auto-generate 11 summaries from Translations/*.md following template format
Format: Red headings + Blue content + CTH style
"""
import sys, io, os, re
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
TRANS_DIR = os.path.join('..', 'Translations')
BANG_TOM_TAT = os.path.join('..', 'BANG_TOM_TAT_11_BAI_BAO.md')

def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)
    return doc

def add_run(doc, text, color=BLUE, bold=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.bold = bold
    return p

def red_bold(doc, text):
    add_run(doc, text, RED, bold=True)

def blue(doc, text, bold=False):
    add_run(doc, text, BLUE, bold)

def red_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RED

def parse_md(filepath):
    """Parse translation MD file into sections"""
    with open(filepath, 'r', encoding='utf-8') as f:
        content = f.read()

    sections = {}
    current = 'header'
    sections[current] = []

    for line in content.split('\n'):
        line = line.strip()
        if not line: continue

        # Detect section headers
        if line.startswith('## TÓM TẮT') or line.startswith('## TOM TAT'):
            current = 'abstract'
            sections[current] = []
        elif line.startswith('## GIỚI THIỆU') or line.startswith('## GI\u1edaI THI\u1ec6U'):
            current = 'intro'
            sections[current] = []
        elif line.startswith('## PHƯƠNG PHÁP') or line.startswith('## PH\u01af\u01a0NG PH\u00c1P'):
            current = 'method'
            sections[current] = []
        elif line.startswith('## KẾT QUẢ') or line.startswith('## K\u1ebeT QU\u1ea2'):
            current = 'results'
            sections[current] = []
        elif line.startswith('## THẢO LUẬN') or line.startswith('## TH\u1ea2O LU\u1eacN'):
            current = 'discussion'
            sections[current] = []
        elif line.startswith('## KẾT LUẬN') or line.startswith('## K\u1ebeT LU\u1eacN'):
            current = 'conclusion'
            sections[current] = []
        elif line.startswith('## TÀI LIỆU') or line.startswith('## REFERENCES'):
            current = 'refs'
            sections[current] = []
        else:
            if current in sections:
                # Skip markdown formatting
                clean = line.lstrip('#').lstrip('- ').lstrip('* ')
                if clean and not clean.startswith('**Tiêu đề') and not clean.startswith('**DOI'):
                    sections[current].append(clean)

    return sections

# Read critique data from BANG_TOM_TAT
critiques = {}
gaps = {}
qualities = {}
with open(BANG_TOM_TAT, 'r', encoding='utf-8') as f:
    bang_content = f.read()

# Parse critique table (Bảng 6)
for line in bang_content.split('\n'):
    if line.startswith('|') and '|' in line[1:]:
        cells = [c.strip() for c in line.split('|')[1:-1]]
        if len(cells) >= 3 and cells[0].isdigit():
            num = int(cells[0])
            if 'Phản biện' in bang_content[bang_content.find(line)-200:bang_content.find(line)]:
                critiques[num] = cells[2] if len(cells) > 2 else ''
            if 'Hướng nghiên cứu' in bang_content[bang_content.find(line)-200:bang_content.find(line)]:
                gaps[num] = cells[2] if len(cells) > 2 else ''

# Paper metadata
meta = {
    1: {'file': '01_Jenkins_et_al_2023', 'quality': '\u2b50\u2b50\u2b50 Trung bình-Khá'},
    2: {'file': '02_Saikia_et_al_2023', 'quality': '\u2b50\u2b50\u2b50 Trung bình'},
    3: {'file': '03_Mandaknalli_Malusare_2021', 'quality': '\u2b50\u2b50 Trung bình-Thấp'},
    4: {'file': '04_NSCH_2020', 'quality': '\u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc'},
    5: {'file': '05_Alharbi_et_al_2019', 'quality': '\u2b50\u2b50\u2b50\u2b50 Khá-Tốt'},
    6: {'file': '06_Nakie_et_al_2022', 'quality': '\u2b50\u2b50\u2b50\u2b50 Tốt'},
    7: {'file': '07_Chen_et_al_2023', 'quality': '\u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc'},
    8: {'file': '08_Wen_et_al_2020', 'quality': '\u2b50\u2b50\u2b50\u2b50 Tốt'},
    9: {'file': '09_Qiu_et_al_2022', 'quality': '\u2b50\u2b50\u2b50\u2b50 Tốt'},
    10: {'file': '10_Xu_et_al_2021', 'quality': '\u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc'},
    11: {'file': '11_Bhardwaj_et_al_2020', 'quality': '\u2b50\u2b50 Trung bình-Thấp'},
}

# Critique data from BANG_TOM_TAT (manually extracted since parsing is complex)
critique_data = {
    1: 'Cỡ mẫu rất nhỏ (N=75), snowball sampling tại 1 trường duy nhất \u2014 không đại diện. GAD-10 sửa đổi thiếu dữ liệu xác thực. Thiết kế kéo dài 2018-2021 trùng COVID-19 nhưng không kiểm soát hiệu ứng thời điểm. Phương pháp hỗn hợp là điểm sáng.',
    2: 'Nam > nữ về lo âu (30,0% vs 18,9%) trái ngược y văn quốc tế \u2014 cần giải thích sâu hơn. DASS-21 tiếng Assam thiếu Cronbach alpha. Chỉ dùng chi-square đơn biến, không hồi quy đa biến. Thu thập 4/2019-6/2020 trùng COVID nhưng không đề cập.',
    3: 'Thiếu nghiêm trọng thông tin phương pháp: công cụ không xác định tên cụ thể, không có Cronbach alpha. Gần 100% có lo âu gợi ý điểm cắt quá thấp. Tạp chí không IF, không PubMed. Không có chấp thuận đạo đức.',
    4: 'Dữ liệu từ cha mẹ, không trực tiếp từ thanh thiếu niên \u2014 có thể underestimate. Dựa trên chẩn đoán y tế \u2014 phụ thuộc tiếp cận dịch vụ. Xu hướng tăng 61% có thể do giảm kỳ thị hơn tăng bệnh thực sự.',
    5: 'Tỷ lệ 74% trầm cảm cực cao do dùng ngưỡng PHQ-9 \u2265 5 (bao gồm nhẹ). So sánh: Chen 2023 23,0%. Chỉ dùng chi-square, không hồi quy đa biến. PHQ-9/GAD-7 chưa xác thực cho bối cảnh Ả Rập Saudi.',
    6: 'Lo âu 66,7% rất cao \u2014 bối cảnh xung đột Ethiopia 2020-2022? AOR hút thuốc CI trên=7304, có thể lỗi in ấn. Phạm vi tuổi rộng (15-25) bao gồm người trưởng thành. DASS-21 tiếng Amharic thiếu xác thực văn hóa.',
    7: 'CI khá rộng dù N=63.205 \u2014 do hiệu ứng thiết kế lấy mẫu cụm. PHQ-9/GAD-7 là sàng lọc, không chẩn đoán. Thiếu mô hình đa tầng phù hợp cấu trúc cụm. Giấc ngủ kém/chơi game có thể là triệu chứng, không phải nguyên nhân.',
    8: 'LPA thiếu báo cáo chỉ số phù hợp mô hình (BIC, AIC, entropy). N=900 từ 1 tỉnh \u2014 hạn chế khái quát hóa. Công cụ đo lo âu không nêu tên cụ thể. "Hỗ trợ tại trường" đo tự đánh giá \u2014 thiên lệch nhận thức.',
    9: 'Giả thuyết trung tâm (tương tác nuôi dạy \u00d7 phục hồi) KHÔNG có ý nghĩa thống kê \u2014 giảm giá trị đóng góp mới. EMBU đo từ quan điểm TTN \u2014 thiên lệch nhận thức tiêu cực. Mẫu 63,4% nam \u2014 mất cân đối giới.',
    10: 'N=373.216 sức mạnh thống kê cao nhưng khác biệt 0,45% đạt P<0,001 mà thiếu ý nghĩa lâm sàng. Thu thập 8 ngày đỉnh COVID \u2014 lo âu tình huống, không phải rối loạn. Biến "worried/fear level" có thể tautology với lo âu.',
    11: 'Tạp chí địa phương, không PubMed, IF không có. Cỡ mẫu nhỏ (n=288). Thiếu phân tích yếu tố liên quan (không hồi quy). Tỷ lệ cực cao cần xem xét ngưỡng cắt DASS-21.',
}

gap_data = {
    1: 'Mở rộng cỡ mẫu. Phát triển công cụ sàng lọc nhạy cảm văn hóa. Nghiên cứu dọc theo dõi từ đầu vị thành niên. Đánh giá can thiệp giải quyết yếu tố văn hóa xã hội.',
    2: 'Nghiên cứu quy mô lớn ở Đông Bắc Ấn Độ. Khảo sát yếu tố văn hóa bộ lạc đặc thù. So sánh thành thị-nông thôn. Giải thích tại sao nam > nữ.',
    3: 'Tái thực hiện với công cụ đo chuẩn hóa (DASS-21 hoặc GAD-7). Nghiên cứu vai trò hoạt động thể chất trong can thiệp lo âu tại trường.',
    4: 'Phân tích nguyên nhân gia tăng 61%. Giải pháp giảm rào cản tiếp cận điều trị. Đánh giá tác động dài hạn COVID-19. Can thiệp cho nhóm thiếu tiếp cận dịch vụ.',
    5: 'Nghiên cứu dọc xu hướng theo thời gian. Vai trò mạng xã hội và áp lực văn hóa Ả Rập. Chương trình SKTT tại trường ở Ả Rập Saudi.',
    6: 'Mở rộng sang khu vực khác Ethiopia và châu Phi. Can thiệp kép: giảm sử dụng chất + hỗ trợ SKTT. Nghiên cứu dọc đánh giá diễn tiến tự nhiên.',
    7: 'So sánh Đông-Tây Trung Quốc. Can thiệp nhắm bắt nạt + giấc ngủ + game. Nghiên cứu dọc đánh giá quỹ đạo triệu chứng.',
    8: 'Áp dụng LPA nhiều vùng nông thôn. Nghiên cứu dọc theo dõi các hồ sơ lo âu. Đánh giá hiệu quả dịch vụ tư vấn tại trường nông thôn.',
    9: 'Nghiên cứu dọc kiểm tra quan hệ nhân quả nuôi dạy \u2192 SKTT. Can thiệp gia đình + tăng cường phục hồi. Nghiên cứu đa văn hóa.',
    10: 'Nghiên cứu dọc hậu COVID-19. So sánh quốc tế cùng cỡ mẫu. Hệ thống cảnh báo sớm SKTT trong khủng hoảng. Cơ chế chênh lệch thành thị-nông thôn.',
    11: 'Tái thực hiện với phân tích mạnh hơn. So sánh trường công-tư. Can thiệp SKTT tại trường công. Khảo sát vai trò kinh tế xã hội chi tiết.',
}

# Generate for each paper
for num in range(1, 12):
    md_file = os.path.join(TRANS_DIR, f'{meta[num]["file"]}.md')
    if not os.path.exists(md_file):
        sys.stderr.write(f'SKIP {md_file} not found\n')
        continue

    sections = parse_md(md_file)
    doc = new_doc()

    # Title
    blue(doc, f'Tóm tắt bài {num}', bold=True)

    # Header from first section
    red_bold(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
    header_lines = sections.get('header', [])
    # Find the main description line (longest non-metadata line)
    header_text = ' '.join(l for l in header_lines if len(l) > 50 and not l.startswith('**'))[:500]
    if header_text:
        blue(doc, header_text)
    else:
        blue(doc, f'(Xem chi tiết trong bản dịch {meta[num]["file"]}.md)')

    # Method
    red_bold(doc, 'Phương pháp nghiên cứu')
    method_lines = sections.get('method', sections.get('abstract', []))
    method_text = ' '.join(method_lines[:5])[:600]
    if method_text:
        blue(doc, method_text, bold=True)

    # Results
    red_bold(doc, 'Kết quả nghiên cứu định lượng')
    result_lines = sections.get('results', [])
    result_text = ' '.join(result_lines[:8])[:800]
    if result_text:
        blue(doc, result_text)
    else:
        # Fallback to abstract
        abstract_lines = sections.get('abstract', [])
        for line in abstract_lines:
            if any(kw in line.lower() for kw in ['%', 'kết quả', 'result', 'tỷ lệ', 'prevalence']):
                blue(doc, line)
                break

    # Discussion/Nhận xét
    red_bold(doc, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
    disc_lines = sections.get('discussion', [])
    disc_text = ' '.join(disc_lines[:5])[:600]
    if disc_text:
        blue(doc, disc_text)

    # Conclusion
    red_bold(doc, 'Kết luận')
    conc_lines = sections.get('conclusion', [])
    conc_text = ' '.join(conc_lines[:4])[:500]
    if conc_text:
        blue(doc, conc_text, bold=True)

    # Critique
    red_h2(doc, 'Phản biện')
    blue(doc, critique_data.get(num, 'Xem bảng tóm tắt'))

    # Gap
    red_h2(doc, 'Hướng nghiên cứu tiếp theo')
    blue(doc, gap_data.get(num, 'Xem bảng tóm tắt'))

    # Quality
    doc.add_paragraph()
    pp = doc.add_paragraph()
    r = pp.add_run(f'Đánh giá chất lượng: {meta[num]["quality"]}')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    fname = f'{meta[num]["file"]}.docx'
    doc.save(fname)
    sys.stderr.write(f'{fname} OK\n')

sys.stderr.write('All 11 done\n')
