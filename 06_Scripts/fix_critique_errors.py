# -*- coding: utf-8 -*-
"""
Fix factual errors detected in the expanded critique:
1. VN001 (Hoa 2024): chỉ Hà Nội — không phải "Hà Nội + Thái Nguyên + Cần Thơ";
   chỉ dùng GAD-7 — không có PHQ-9; 40,6% ở GAD-7 ≥5 — không phải 23,5% ở ≥10
2. VN029 (Duong 2025): cắt ngang đa trung tâm tại TPHCM (n=2.631), dùng DASS-21 + YBRS
   — không phải longitudinal 3 timepoints; không có "persistence rate 55%"
3. VN030 (Happy House): tại Hà Nội — không phải TPHCM; CCT (cluster controlled
   non-randomized), n=1.084, CESD-R là endpoint chính; d=0,11 universal, fade-out 6 tháng
4. Citations ngoài refs list (Creswell&Plano-Clark 2017, Goodman 1997, Rutter 2007,
   Meltzer 2000, Tong 2007): đánh dấu rõ là "khung phương pháp bên ngoài"
5. Dang 2017 Cronbach α cụ thể: chưa verify — đổi thành câu chung
6. Hirota 2020 ratio tư vấn viên Thailand/VN: chưa verify — làm mềm
7. Cong 2018 30–40% victim cyberbullying: chưa verify — làm mềm
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

FIXES = [
    # 1. VN001 errors
    ('VN001 (Hoa et al. 2024, n = 3.910, Hà Nội + Thái Nguyên + Cần Thơ)',
     'VN001 (Hoa et al. 2024, n = 3.910 HS THPT, chỉ Hà Nội)'),
    ('Dùng GAD-7 và PHQ-9 — không trực tiếp so được SDQ, nhưng cùng chỉ ra nữ > nam (OR ≈ 2,0 cho GAD-7 cut-off 10), THPT > THCS. Bổ sung cho UNICEF 2022: prevalence lo âu tổng quát (GAD-7 ≥ 10) 23,5 %, gần với 26,1 % của UNICEF.',
     'Chỉ dùng GAD-7 (không có PHQ-9) với phỏng vấn định tính. Bổ sung cho UNICEF 2022: prevalence lo âu GAD-7 ≥ 5 là 40,6 %, nữ có điểm GAD-7 cao hơn nam (M = 1,74 vs 1,50, p < 0,01) — cùng xu hướng giới với UNICEF 2022. Lưu ý: GAD-7 ≥ 5 là sàng lọc không phải chẩn đoán; so sánh với SDQ-25 cần cẩn trọng vì đo construct khác nhau.'),

    # 2. VN029 errors
    ('VN029 (Duong et al. 2025, n = 2.631, longitudinal 3 timepoints)',
     'VN029 (Thai, Duong et al. 2025, n = 2.631, cắt ngang đa trung tâm TPHCM)'),
    ('Thiết kế longitudinal khắc phục hạn chế cắt ngang của UNICEF 2022. Báo cáo tracking quỹ đạo triệu chứng qua 2 năm; cho thấy HS có nguy cơ cao ban đầu (SDQ-like) duy trì nguy cơ cao ở timepoint 2 (persistence rate ≈ 55 %).',
     'Khảo sát cắt ngang 8 cơ sở giáo dục TPHCM (4 THPT chính quy + 4 GDTX) dùng DASS-21 (trầm cảm/lo âu/căng thẳng) + YBRS (8 nhóm hành vi nguy cơ). Tỷ lệ lo âu 50,3 % — cao hơn SDQ emotional subscale của UNICEF 2022 (30,9 %), phản ánh DASS-21 đo broader range còn SDQ đo emotional problems hẹp hơn. Mối liên hệ: HS có triệu chứng SKTT có nguy cơ tham gia hành vi nguy cơ cao hơn 1,24–4,64 lần — không có trong UNICEF 2022 và là đóng góp riêng. Tuy nhiên cả UNICEF 2022 lẫn VN029 đều là cắt ngang — chưa có nghiên cứu longitudinal toàn quốc tại VN.'),

    # 3. VN030 errors
    ('Đây là can thiệp MHPSS trường học có đối chứng tại TPHCM — bổ sung nhu cầu "RCT can thiệp MHPSS học đường" mà UNICEF 2022 xác định là gap. Tham khảo cũng liên kết citation #901 trong refs của UNICEF: Dang, Weiss, Nguyen, Tran & Pollack 2016 về efficacy can thiệp RISE tại VN.',
     'Cluster controlled trial (non-randomized) tại 8 trường Hà Nội (4 can thiệp + 4 đối chứng), n = 1.084 HS lớp 10. Can thiệp Happy House (RAP-A adapt văn hoá VN) do giáo viên Giáo dục Công dân cung cấp 6 buổi × 90 phút. Endpoint chính CESD-R ≥ 16; effect size d = 0,11 cho universal intervention, fade-out ở 6 tháng nhưng CSES ứng phó vẫn duy trì. Đây là CCT ĐẦU TIÊN tại VN cho MHPSS trường học — lấp khoảng trống mà UNICEF 2022 chỉ ra. Lưu ý: universal → hiệu ứng nhỏ; khuyến nghị hướng targeted (HS có GAD-7 ≥ 8 hoặc CESD-R ≥ 16). Liên kết với UNICEF refs #901 (Dang, Weiss, Nguyen, Tran & Pollack 2016 — can thiệp RISE tại VN).'),

    # 4. External citations flagging
    ('Creswell & Plano Clark (2017) xem là "integrative mixed-methods"',
     'Creswell & Plano Clark (2017, khung bên ngoài refs UNICEF) xem là "integrative mixed-methods"'),
    ('Goodman 1997; validate VN bởi Dang et al. 2016',
     'Goodman 1997 — khung bên ngoài; validate VN bởi Dang et al. 2016'),
    ('(Rutter 2007, "Proceeding from observed correlation to causal inference: The use of natural experiments")',
     '(Rutter 2007 — khung bên ngoài refs UNICEF: "Proceeding from observed correlation to causal inference")'),
    ('Meltzer et al. 2000 (British Child Mental Health Survey) cho thấy teacher rating và self-report của trẻ em có correlation chỉ 0,20–0,35',
     'Meltzer et al. 2000 (khung bên ngoài — British Child Mental Health Survey) cho thấy teacher rating và self-report có correlation thấp, điển hình 0,20–0,35'),
    ('COREQ (Consolidated Criteria for Reporting Qualitative Research — Tong, Sainsbury & Craig 2007)',
     'COREQ (Consolidated Criteria for Reporting Qualitative Research — Tong, Sainsbury & Craig 2007, khung bên ngoài)'),

    # 5. Dang 2017 alpha specific — soften
    ('Dang et al. 2017 (tham chiếu citation trong báo cáo) đã chỉ ra factor structure 5-subscale của SDQ-25 không ổn định trong mẫu Việt Nam — Cronbach α của subscale "conduct problems" và "peer problems" dao động 0,47–0,62 (dưới ngưỡng chấp nhận 0,70). Báo cáo UNICEF không công bố reliability riêng cho mẫu của mình (trang 35 chỉ tham chiếu "đã validate"). Đây là gap cần minh bạch trong báo cáo phiên bản tiếp theo.',
     'Dang và cộng sự (2017, có citation trong báo cáo) đã đặt câu hỏi về factor structure 5-subscale của SDQ-25 trong bối cảnh Việt Nam; một số subscale (đặc biệt conduct và peer problems) có reliability không đạt ngưỡng chấp nhận 0,70 trong mẫu châu Á. Báo cáo UNICEF không công bố reliability riêng (Cronbach α) cho mẫu 668 HS của mình — chỉ nêu "đã validate tiếng Việt". Đây là gap minh bạch cần bổ sung trong báo cáo wave sau (alpha cụ thể cho từng subscale).'),

    # 6. Hirota 2020 Thailand/VN ratio — soften (I don't have exact numbers)
    ('review của Hirota, Guerrero & Skokauskas 2020 (citation #919) chỉ ra tỷ lệ tư vấn viên học đường / HS ở Thailand ~ 1:1.200 — cao hơn đáng kể so với VN ~ 1:4.500 (theo khuyến nghị MOET).',
     'review của Hirota, Guerrero & Skokauskas 2020 (citation #919) cho thấy Thailand có tỷ lệ school counselor / HS thuận lợi hơn đáng kể so với VN; con số cụ thể phụ thuộc thời điểm và nguồn (cần verify trực tiếp từ paper).'),

    # 7. Cong 2018 cyberbullying 30-40% — soften
    ('Cong et al. 2018 (n = 770 HS Hà Nội) đo tỷ lệ victim 30–40 %; UNICEF 2022 cao hơn, có thể do đo broader exposure (hiếm + đôi khi + thường) và giai đoạn COVID tăng online time',
     'Cong et al. 2018 (VN HS Hà Nội, citation #898) là validation RCBI tiếng Việt; UNICEF 2022 dùng công cụ này và ghi nhận 52,2 % bị ảnh hưởng cyberbullying — cao hơn một số ước tính pre-COVID cho VN, phản ánh (a) broader exposure categorisation (hiếm + đôi khi + thường) và (b) gia tăng online time thời COVID'),
]

d = Document(PATH)
changes = 0
for p in d.paragraphs:
    t = p.text
    if not t.strip():
        continue
    new_t = t
    for find, rep in FIXES:
        if find in new_t:
            new_t = new_t.replace(find, rep)
    if new_t != t:
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = new_t
        changes += 1

d.save(PATH)
print(f'Fixed {changes} paragraphs')

# Verify no remaining errors
d2 = Document(PATH)
txt = '\n'.join(p.text for p in d2.paragraphs)
print('\nVerification:')
checks = [
    ('Hà Nội + Thái Nguyên + Cần Thơ', 0),  # should be 0
    ('longitudinal 3 timepoints', 0),  # should be 0
    ('persistence rate', 0),  # should be 0
    ('MHPSS trường học có đối chứng tại TPHCM', 0),  # should be 0
    ('chỉ Hà Nội', 1),  # should appear
    ('cắt ngang đa trung tâm TPHCM', 1),  # should appear
    ('tại 8 trường Hà Nội', 1),  # should appear
    ('khung bên ngoài', 5),  # should appear multiple times
]
for phrase, expected in checks:
    actual = txt.count(phrase)
    if expected > 0:
        status = 'OK' if actual >= expected else 'FAIL'
    else:
        status = 'OK' if actual == 0 else 'FAIL'
    print(f'  "{phrase}" appears {actual} times (expected {">=" if expected > 0 else "="} {expected}): {status}')
