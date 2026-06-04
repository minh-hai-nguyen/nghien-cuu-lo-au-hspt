# -*- coding: utf-8 -*-
"""Ban tom tat LA tieng Anh - 24 trang.
NCS Cong Thi Hang - 29/05/2026.
Format A5, TNR 11pt, theo chuan TS tai VN."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING, WD_BREAK
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', 'TomTatLA_EN_v1_29052026.docx')


def doc_init():
    d = Document()
    s = d.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(11)
    s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    s.paragraph_format.line_spacing = 1.15
    s.paragraph_format.space_after = Pt(3)
    for sec in d.sections:
        sec.page_width = Cm(14.8); sec.page_height = Cm(21.0)
        sec.top_margin = Cm(1.5); sec.bottom_margin = Cm(1.5)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.5)
    return d


def H1(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True

def H2(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def H3(d, text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True

def P(d, text, indent=True, bold=False, italic=False, center=False, size=11):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER if center else WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent and not center:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic

def page_break(d):
    p = d.add_paragraph()
    r = p.add_run()
    r.add_break(WD_BREAK.PAGE)

def add_table_simple(d, headers, rows, widths_cm=None):
    t = d.add_table(rows=1, cols=len(headers))
    t.style = 'Light Grid Accent 1'
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        if widths_cm: hdr[i].width = Cm(widths_cm[i])
    for row_data in rows:
        row = t.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            r = p.add_run(str(cell_data))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
            if widths_cm: row[i].width = Cm(widths_cm[i])


# ============================================================
d = doc_init()

# COVER
P(d, 'MINISTRY OF EDUCATION AND TRAINING', center=True, bold=True, size=12, indent=False)
P(d, 'HANOI NATIONAL UNIVERSITY OF EDUCATION', center=True, bold=True, size=12, indent=False)
d.add_paragraph(); d.add_paragraph(); d.add_paragraph()
P(d, 'CONG THI HANG', center=True, bold=True, size=14, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'ANXIETY DISORDERS AMONG', center=True, bold=True, size=16, indent=False)
P(d, 'LOWER SECONDARY SCHOOL STUDENTS', center=True, bold=True, size=16, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'Specialty: Specialised Psychology', center=True, size=12, indent=False)
P(d, 'Code: 9.31.04.01', center=True, size=12, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'SUMMARY OF DOCTORAL DISSERTATION IN PSYCHOLOGY', center=True, bold=True, size=13, indent=False)
d.add_paragraph(); d.add_paragraph(); d.add_paragraph()
P(d, 'HANOI, 2026', center=True, bold=True, size=12, indent=False)
page_break(d)

# INNER COVER
P(d, 'The dissertation was completed at:', bold=True, indent=False, size=11)
P(d, 'HANOI NATIONAL UNIVERSITY OF EDUCATION', bold=True, center=True, size=12, indent=False)
d.add_paragraph()
P(d, 'Scientific supervisor:', bold=True, indent=False, size=11)
P(d, 'Dr. Dao Minh Duc', center=True, italic=True, size=11, indent=False)
d.add_paragraph()
P(d, 'Reviewer 1: .........................................................................................', indent=False, size=11)
P(d, '            .........................................................................................', indent=False, size=11)
d.add_paragraph()
P(d, 'Reviewer 2: .........................................................................................', indent=False, size=11)
P(d, '            .........................................................................................', indent=False, size=11)
d.add_paragraph()
P(d, 'Reviewer 3: .........................................................................................', indent=False, size=11)
P(d, '            .........................................................................................', indent=False, size=11)
d.add_paragraph(); d.add_paragraph()
P(d, 'The dissertation will be defended before the University-level Dissertation Evaluation Committee at Hanoi National University of Education', size=11)
P(d, 'at ........ on ........ / ........ / ........', size=11, indent=False)
d.add_paragraph(); d.add_paragraph()
P(d, 'The dissertation can be consulted at:', bold=True, indent=False, size=11)
P(d, '- National Library of Vietnam', size=11, indent=False)
P(d, '- Library of Hanoi National University of Education', size=11, indent=False)
page_break(d)

# ========================================================
# INTRODUCTION
# ========================================================
H1(d, 'INTRODUCTION')

H2(d, '1. Rationale for the study')
P(d, 'In the current globally turbulent context, the mental health of adolescents is facing unprecedented challenges, among which Anxiety Disorders (AD) have emerged as one of the most prevalent and prominent issues. The World Health Organization (WHO) emphasises that anxiety disorders are among the two most serious mental health concerns, with early onset typically occurring during adolescence. Overall prevalence increased by up to 25% during the COVID-19 pandemic.')
P(d, 'In Vietnam, the nationally representative V-NAMHS survey (Institute of Sociology et al., 2022) of 5,996 parent-adolescent dyads revealed that anxiety disorders are the most prevalent mental disorders among 10–17-year-olds, with a DSM-5 diagnostic rate of 2.3%. However, screening-based studies have reported widely varying rates, sometimes exceeding 50%, indicating substantial methodological gaps in the measurement of anxiety disorders among Vietnamese students.')
P(d, 'Lower secondary school students (aged 11–15) occupy a particularly important developmental stage, characterised by vulnerability to mental health challenges due to the complex interaction between psychophysiological changes and academic, familial, and social pressures. This is also the typical onset stage for several types of anxiety disorders, particularly social anxiety disorder.')
P(d, 'Existing domestic research presents five major gaps: (1) lack of a standardised anxiety disorder assessment tool culturally and linguistically adapted for Vietnamese children and adolescents; (2) most studies focus on generalised anxiety, while social anxiety and separation anxiety remain largely unexplored; (3) protective factors have not been examined simultaneously and integrated with risk factors in a unified design; (4) longitudinal studies remain rare; (5) limited research on cognitive-behavioural therapy (CBT) interventions for Vietnamese lower secondary school students.')
P(d, 'Based on these gaps, the dissertation "Anxiety Disorders among Lower Secondary School Students" was selected to provide empirical evidence on the prevalence and contributing factors, while proposing a prevention training framework suited to the Vietnamese educational context.')

H2(d, '2. Research objectives')
P(d, 'To investigate the theoretical foundations of anxiety disorders among lower secondary school students; to examine the prevalence of anxiety disorders and contributing factors (both risk and protective); and on that basis, to propose a training framework for the prevention of anxiety disorders among lower secondary school students.')

H2(d, '3. Research subjects')
P(d, 'The levels, manifestations, and impacts of factors influencing school-based anxiety disorders among lower secondary school students; the proposed framework for anxiety prevention training for these students.')

H2(d, '4. Research participants')
P(d, '- Lower secondary school students: a survey of 1,352 students and in-depth interviews with 12 students in Hanoi.', indent=False)
P(d, '- Lower secondary school teachers: in-depth interviews with 6 homeroom teachers.', indent=False)

H2(d, '5. Scientific hypotheses')
P(d, 'The dissertation formulates three scientific hypotheses:')
P(d, 'Hypothesis 1 — Gender differences depend on the type of anxiety disorder: for generalised anxiety and social anxiety, female students show higher levels than male students; for separation anxiety, there is no statistically significant gender difference.')
P(d, 'Hypothesis 2 — Self-esteem is among the strongest protective factors, with effect size comparable to that of academic pressure on anxiety disorders.')
P(d, 'Hypothesis 3 — Academic pressure and mobile phone addiction are among the strongest risk factors influencing anxiety disorders in lower secondary school students.')

H2(d, '6. Research tasks')
P(d, '(1) To systematise the theoretical foundations of anxiety disorders and anxiety disorders among lower secondary school students.', indent=False)
P(d, '(2) To investigate the prevalence, manifestations, and contributing factors (risk factors, protective factors, and coping strategies) of anxiety disorders among lower secondary school students.', indent=False)
P(d, '(3) To propose a school-based training framework for the prevention of anxiety disorders among lower secondary school students.', indent=False)

H2(d, '7. Scope of the study')
P(d, 'Content: focuses on three types of anxiety disorders (generalised, social, separation) and contributing factors (academic pressure, mobile phone addiction, school bullying, self-esteem, school attachment, social support, and coping strategies).')
P(d, 'Setting: 2 lower secondary schools in Hanoi — Nhat Tan Lower Secondary School (inner city) and Tay Mo Lower Secondary School (peri-urban).')
P(d, 'Participants: lower secondary school students from Grades 6–9 (aged 11–14).')

H2(d, '8. Research methods')
P(d, 'The dissertation employs a combination of methods: (1) literature and document review; (2) questionnaire survey; (3) in-depth interview; (4) observation; (5) statistical analysis using SPSS 31.0 and AMOS 31.0.')

H2(d, '9. New contributions of the dissertation')
P(d, 'Theoretical contributions: systematising the theoretical foundations of anxiety disorders among lower secondary school students by integrating networks of risk and protective factors; establishing a conceptual framework for coping strategies along two dimensions — adaptive and maladaptive.')
P(d, 'Practical contributions: providing quantitative evidence on the prevalence of anxiety disorders among 1,352 lower secondary school students in Hanoi; proposing a scientifically grounded and practically feasible training framework for the prevention of school-based anxiety disorders, applicable to Vietnamese school settings.')

H2(d, '10. Structure of the dissertation')
P(d, 'In addition to the Introduction, Conclusion, References, and Appendices, the dissertation comprises 3 chapters: Chapter 1 — Theoretical foundations of anxiety disorders among lower secondary school students; Chapter 2 — Research organisation and methods; Chapter 3 — Empirical results on the prevalence of anxiety disorders.')

page_break(d)

# CHAPTER 1
H1(d, 'CHAPTER 1\nTHEORETICAL FOUNDATIONS OF ANXIETY\nDISORDERS AMONG LOWER SECONDARY SCHOOL STUDENTS')

H2(d, '1.1. Literature review')
H3(d, '1.1.1. Prevalence and levels of anxiety disorders')
P(d, 'International studies indicate wide variation in anxiety disorder prevalence among adolescents, depending on the screening instrument and national context. Xu et al. (2021), surveying 373,216 secondary school students in China using the GAD-7, reported a prevalence of 9.89% at moderate or higher levels, rising to 38.42% when including the mild level. Bhardwaj et al. (2020), surveying 288 upper secondary students in India using the DASS-21, found prevalences of 81.9% (mild or higher) and 73.2% (moderate or higher). Chen et al. (2023), surveying 63,205 secondary school students in China, reported anxiety at 13.9% and depression at 23.0%.')
P(d, 'In Vietnam, the V-NAMHS survey (Institute of Sociology et al., 2022) of 5,996 parent-adolescent dyads identified an anxiety disorder prevalence meeting DSM-5 criteria of 2.3%. Other screening-based studies reported higher rates: Hoang Trung Hoc and Nguyen Thuy Dung (2025), surveying 8,473 students across six provinces, reported anxiety rates of 41.5% during COVID-19 and 25.4% post-COVID-19; Nguyen Ngoc Bao Quyen et al. (2025), surveying 501 upper secondary students in Hanoi using the DASS-21, reported a rate of 86.2%.')

H3(d, '1.1.2. Risk and protective factors')
P(d, 'Key risk factors identified in international literature include: academic pressure (Steare et al., 2023 — a systematic review of 52 studies, 48 of which confirmed a positive correlation between academic pressure and poor mental health); mobile phone and social media addiction (A.M. Saikia et al., 2023; Z. Chen et al., 2023; T.L. Anderson et al., 2025); school bullying (Z. Chen et al., 2023 — reporting school bullying rates ranging from 9.0% to 61.6%).')
P(d, 'Protective factors include: self-esteem (Rosenberg, 1965; Masten, 2014); school attachment; parental and peer support (Compas et al., 2017 — a meta-analysis of 212 studies, N = 80,850); and adaptive coping strategies (Lazarus and Folkman, 1984; Compas et al., 2017).')

H3(d, '1.1.3. Research gaps in Vietnam')
P(d, 'Review of domestic studies reveals five key gaps: (1) lack of a standardised anxiety disorder assessment tool culturally and linguistically adapted for Vietnamese children and adolescents, resulting in widely varying prevalence estimates (from 2.3% by full diagnosis to over 50% by screening); (2) most studies focus on generalised anxiety using the GAD-7/DASS-21, while social and separation anxiety remain largely unstudied; (3) protective factors have not been examined simultaneously and integrated with risk factors in a unified design; (4) longitudinal studies and multi-centre randomised clinical trials are rare; (5) limited research on CBT-based interventions for Vietnamese lower secondary school students.')

H2(d, '1.2. Theoretical foundations')
H3(d, '1.2.1. Concept of anxiety disorders')
P(d, 'Anxiety disorder is an abnormal psychological state involving excessive fear, persistent worry, and related behavioural disturbances that cause significant distress in personal life. According to DSM-5 (American Psychiatric Association, 2013), anxiety disorders are structured as a multi-component system, including cognitive aspects (negative thinking, catastrophising), emotional aspects (tension, restlessness), physiological aspects (hyperactivation of the sympathetic nervous system), and behavioural aspects (avoidance, repetitive checking).')

H3(d, '1.2.2. Classification of anxiety disorders according to DSM-5')
P(d, 'According to DSM-5, the main forms of anxiety disorders in children and adolescents include: (1) Generalised Anxiety Disorder (GAD): excessive worry about multiple events lasting at least 6 months, accompanied by somatic symptoms; (2) Social Anxiety Disorder: marked and persistent fear in social or performance situations; (3) Separation Anxiety Disorder: excessive concern about separation from attachment figures, typically with onset in childhood and adolescence.')

H3(d, '1.2.3. Lower secondary school students and anxiety disorders')
P(d, 'Anxiety disorders among lower secondary school students are abnormal psychological states characterised by emotional imbalance, including excessive fear, worry, disproportionate fear responses to the perceived danger of a situation, and related behavioural disturbances, accompanied by physiological manifestations, persistent and detrimental to academic performance and quality of life.')

H3(d, '1.2.4. Network of contributing factors')
P(d, 'The dissertation establishes an analytical framework for factors influencing anxiety disorders among lower secondary school students, organised into two groups: risk factors (academic pressure, mobile phone addiction, physical bullying, verbal and cyber bullying, low self-esteem) and protective factors (self-esteem, school attachment, parental support, peer support, adaptive coping strategies). These factors interact simultaneously and are not mutually exclusive.')

H2(d, '1.3. Training framework for the prevention of anxiety disorders')
P(d, 'Based on the findings regarding risk and protective factors, the dissertation develops a training framework for the prevention of anxiety disorders among lower secondary school students, with three components: (1) scientific foundations based on the cognitive-behavioural therapy (CBT) approach and the stepped care model; (2) framework concept — a programme oriented towards primary and secondary prevention rather than clinical treatment; (3) programme structure comprising eight core modules, with implementation procedures and effectiveness evaluation.')

page_break(d)

# CHAPTER 2
H1(d, 'CHAPTER 2\nRESEARCH ORGANISATION AND METHODS')

H2(d, '2.1. Research organisation')
P(d, 'The study employs a mixed quantitative–qualitative design, implemented in two phases:')
P(d, 'Theoretical phase: literature synthesis, conceptual framework development, and selection and adaptation of measurement instruments suited to Vietnamese lower secondary school students.')
P(d, 'Empirical phase: large-scale questionnaire surveys and in-depth interviews to supplement quantitative findings.')
P(d, 'Research settings: 2 lower secondary schools in Hanoi — Nhat Tan Lower Secondary School (inner city) and Tay Mo Lower Secondary School (peri-urban).')

H2(d, '2.2. Participants')
P(d, 'Quantitative sample: 1,352 lower secondary school students from Grades 6–9 (aged 11–14), selected through stratified sampling by school and grade level, combined with cluster sampling at the classroom level. The sample comprised 614 male students (45.4%) and 738 female students (54.6%).')
P(d, 'Qualitative sample: 12 lower secondary school students for in-depth interviews; 6 homeroom teachers for in-depth interviews.')

H2(d, '2.3. Research methods')
H3(d, '2.3.1. Questionnaire survey')
P(d, 'The questionnaire comprises two parts: Part 1 gathers general information about participants (gender, grade, academic performance, health, psychological, and behavioural characteristics); Part 2 assesses the prevalence of anxiety disorders and contributing factors using internationally validated and adapted measurement scales.')

H3(d, '2.3.2. Measurement instruments')
P(d, 'Measurement scales were selected based on international validation and adapted through forward–back translation, expert consultation, and pilot testing:')
P(d, '- Anxiety disorder measurement: RCADS (Revised Children\'s Anxiety and Depression Scale; Chorpita, 2000), differentiating three forms — generalised anxiety disorder (GAD), separation anxiety disorder (SAD), and social anxiety disorder (SocAD).', indent=False)
P(d, '- Risk factor measurement: ESSA (Educational Stress Scale for Adolescents; Sun et al., 2011) for academic pressure; SAS-SV (Smartphone Addiction Scale - Short Version; Kwon et al., 2013) for smartphone addiction; OBVQ (Olweus Bully/Victim Questionnaire; Olweus, 1996) for school bullying, with two factors: physical bullying and verbal bullying.', indent=False)
P(d, '- Protective factor measurement: PSSM (Psychological Sense of School Membership; Goodenow, 1993) for school attachment; MSPSS (Multidimensional Scale of Perceived Social Support; Zimet et al., 1988) for perceived social support, with separate parental and peer support subscales; Brief COPE (Carver, 1997) for coping strategies, with three factors: avoidance, problem-solving coping, and support-seeking; RSES (Rosenberg Self-Esteem Scale; Rosenberg, 1965) for self-esteem.', indent=False)
P(d, 'All adapted scales achieved acceptable reliability (Cronbach\'s α ≥ 0.70 and McDonald\'s ω ≥ 0.70) and structural validity via Confirmatory Factor Analysis (CFA) with fit indices reaching acceptable thresholds.')

H3(d, '2.3.3. In-depth interview and observation methods')
P(d, 'In-depth interviews aimed to collect qualitative data clarifying further the manifestations and experiences related to anxiety disorders. Observation captured behavioural manifestations, attitudes, and reactions of students during the survey process.')

H3(d, '2.3.4. Statistical analysis methods')
P(d, 'Data were processed using SPSS 31.0 (data cleaning, reliability assessment via Cronbach\'s α and McDonald\'s ω, descriptive statistics, t-tests, ANOVA) and AMOS 31.0 (Confirmatory Factor Analysis — CFA, Structural Equation Modelling — SEM).')

H2(d, '2.4. Research ethics')
P(d, 'The study adhered to ethical principles for research involving children and adolescents: (1) informed consent from parents and assent from students; (2) confidentiality and data protection; (3) the right to withdraw at any time; (4) provision of psychological counselling referral for students identified with high-level anxiety symptoms during the survey.')

page_break(d)

# CHAPTER 3
H1(d, 'CHAPTER 3\nEMPIRICAL RESULTS ON THE PREVALENCE\nOF ANXIETY DISORDERS')

H2(d, '3.1. Measurement validity of the scales')
P(d, 'Analysis showed that all scales used in the study achieved good reliability and acceptable structural fit on the Vietnamese lower secondary school student sample.')

P(d, 'Table 3.1. Summary of reliability and CFA fit indices of the scales', bold=True, italic=True, center=True, indent=False, size=10)
add_table_simple(d,
    headers=['Scale', 'α', 'ω', 'CFI', 'RMSEA'],
    rows=[
        ['Generalised anxiety (GAD)', '0.811', '0.811', '0.982', '0.049'],
        ['Separation anxiety (SAD)', '0.726', '0.726', '0.975', '0.099'],
        ['Social anxiety (SocAD)', '0.744', '0.750', '0.992', '0.060'],
        ['Academic pressure (AP)', '0.708', '0.716', '0.998', '0.024'],
        ['Mobile phone addiction (MPA)', '0.836', '0.839', '0.996', '0.039'],
        ['Physical bullying', '0.775', '0.777', '0.995', '0.050'],
        ['Verbal/cyber bullying', '0.864', '0.864', '0.995', '0.072'],
        ['School attachment', '0.747', '0.746', '0.978', '0.042'],
        ['Parental support', '0.847', '0.848', '0.996', '0.061'],
        ['Peer support', '0.837', '0.837', '0.995', '0.064'],
        ['Avoidance (coping)', '0.727', '0.705', '0.862', '0.169'],
        ['Problem-solving coping', '0.700', '0.699', '0.972', '0.052'],
        ['Support-seeking', '0.773', '0.775', '1.000', '0.000'],
        ['Self-esteem', '0.725', '0.724', '0.988', '0.045'],
    ],
    widths_cm=[5.5, 1.5, 1.5, 1.5, 1.5])
P(d, 'All scales met the reliability thresholds (α ≥ 0.70; ω ≥ 0.70), with most CFA models achieving good fit (CFI ≥ 0.90; RMSEA ≤ 0.10), thereby confirming the measurement validity of the instruments.', italic=True, size=10)

H2(d, '3.2. Prevalence of anxiety disorders among lower secondary school students')
P(d, 'Survey results on 1,352 lower secondary school students indicated that anxiety disorders are primarily at moderate levels; generalised anxiety tends to be higher than the other forms, social anxiety is at moderate level, and separation anxiety is at low level. Anxiety manifestations are concentrated mainly in academic pressure, self-evaluation, and social situations.')

P(d, 'Table 3.2. Mean scores of anxiety disorders by gender (0–100 scale)', bold=True, italic=True, center=True, indent=False, size=10)
add_table_simple(d,
    headers=['Anxiety type', 'Male (n=614)\nM (SD)', 'Female (n=738)\nM (SD)', 'F', 'p'],
    rows=[
        ['Generalised anxiety', '51.43 (22.01)', '59.47 (22.07)', '44.484', '< 0.001'],
        ['Separation anxiety', '25.42 (25.46)', '24.76 (23.29)', '0.246', '0.620'],
        ['Social anxiety', '43.20 (25.09)', '52.74 (26.31)', '45.984', '< 0.001'],
        ['Total anxiety', '40.02 (19.02)', '45.66 (18.91)', '29.642', '< 0.001'],
    ],
    widths_cm=[3.2, 2.7, 2.7, 1.4, 1.4])
P(d, 'The results confirm Hypothesis 1: female students show significantly higher levels of generalised anxiety and social anxiety than male students (p < 0.001), whereas separation anxiety shows no gender difference (F = 0.246; p = 0.620).', italic=True, size=10)

H2(d, '3.3. Prevalence of risk and protective factors')
P(d, 'Descriptive analysis indicated mean scores (converted to a 0–100 scale) as shown in Table 3.3.')

P(d, 'Table 3.3. Mean scores of risk and protective factors (0–100 scale)', bold=True, italic=True, center=True, indent=False, size=10)
add_table_simple(d,
    headers=['Group', 'Factor', 'M', 'SD'],
    rows=[
        ['Risk', 'Academic pressure', '51.13', '23.92'],
        ['Risk', 'Mobile phone addiction', '28.38', '—'],
        ['Risk', 'Physical bullying', '13.52', '—'],
        ['Protective', 'School attachment', '52.60', '20.02'],
    ],
    widths_cm=[2.0, 5.0, 1.8, 1.8])
P(d, 'Among risk factors, academic pressure exceeds the scale midpoint (51.13/100), significantly higher than mobile phone addiction (28.38) and physical bullying (13.52).', italic=True, size=10)

H2(d, '3.4. Impact of contributing factors on anxiety disorders')
H3(d, '3.4.1. Impact of risk factors')
P(d, 'Structural Equation Modelling (SEM) results revealed three groups of risk factors ranked by strength of impact on total anxiety disorders:')

P(d, 'Table 3.4. Standardised β coefficients: Risk factors → Anxiety disorders', bold=True, italic=True, center=True, indent=False, size=10)
add_table_simple(d,
    headers=['Risk factor', '→ GAD', '→ SAD', '→ SocAD', '→ Total AD (3 factors)'],
    rows=[
        ['Academic pressure (AP)', '0.510***', '0.253***', '0.490***', '0.533*** (R²=0.284)'],
        ['Mobile phone addiction (MPA)', '0.336***', '0.265***', '0.383***', '0.400*** (R²=0.160)'],
        ['School bullying (SB)', '0.215***', '0.376***', '0.253***', '0.276*** (R²=0.076)'],
    ],
    widths_cm=[3.5, 1.7, 1.7, 1.7, 3.0])
P(d, 'Note: *** p < 0.001. AP is the strongest risk factor (β = 0.533). MPA ranks second (β = 0.400). SB shows a distinctive pattern: its strongest impact is on separation anxiety (β = 0.376), rather than generalised or social anxiety.', italic=True, size=10)
P(d, 'These results confirm Hypothesis 3: academic pressure and mobile phone addiction are among the strongest risk factors.', italic=True, size=10)

H3(d, '3.4.2. Impact of protective factors')
P(d, 'Analysis showed that protective factors exert negative effects (reducing anxiety) of varying strength:')

P(d, 'Table 3.5. Standardised β coefficients: Protective factors → Anxiety disorders', bold=True, italic=True, center=True, indent=False, size=10)
add_table_simple(d,
    headers=['Protective factor', '→ GAD', '→ SAD', '→ SocAD', '→ Total AD (3 factors)'],
    rows=[
        ['Self-esteem (SE)', '-0.455***', '-0.087*', '-0.415***', '-0.457*** (R²=0.209)'],
        ['Parental support (PS)', '-0.172***', '0.000', '-0.273***', '-0.234*** (R²=0.055)'],
        ['School attachment (SA)', '-0.108**', '0.014', '-0.187***', '-0.155*** (R²=0.024)'],
    ],
    widths_cm=[3.5, 1.7, 1.7, 1.7, 3.0])
P(d, 'Note: *** p < 0.001; ** p < 0.01; * p < 0.05. Self-esteem is the strongest protective factor (β = -0.457), with effect strength approximately 85-89% of academic pressure (per LA Conclusions, depending on specific anxiety disorder type).', italic=True, size=10)
P(d, 'These results confirm Hypothesis 2: self-esteem is among the strongest protective factors, with effect strength comparable to academic pressure on generalised anxiety and social anxiety.', italic=True, size=10)

H3(d, '3.4.3. Integrated risk and protective factor model')
P(d, 'The integrated model combining risk and protective factors revealed that risk factors dominate protective factors at the aggregate level, suggesting the need for concurrent intervention across both dimensions: reducing risk while strengthening protective factors.')

H2(d, '3.5. Coping strategies of lower secondary school students')
P(d, 'Analysis of the Brief COPE along three factors (avoidance, problem-solving coping, support-seeking) indicated:')
P(d, '- Maladaptive dimension (avoidance): increases markedly alongside anxiety symptoms — consistent with theoretical expectations.', indent=False)
P(d, '- Adaptive dimension (problem-solving and support-seeking): also correlates positively with anxiety symptoms — contrary to expectations, suggesting that students are consciously using positive strategies but have not yet achieved the expected effectiveness.', indent=False)
P(d, 'These findings suggest an intervention direction: simultaneously (1) reducing maladaptive strategies and (2) enhancing the quality of implementation of adaptive strategies.')

H2(d, '3.6. Proposed training framework for the prevention of anxiety disorders')
P(d, 'Based on the theoretical and empirical findings, the dissertation proposes a Training Framework for the Prevention of Anxiety Disorders among Lower Secondary School Students, comprising eight core modules: (1) Recognising anxiety disorders and self-assessment; (2) Cognitive restructuring skills; (3) Academic pressure management skills; (4) Healthy digital device use skills; (5) Coping skills for school bullying; (6) Strengthening self-esteem; (7) Enhancing school attachment and social support; (8) Developing adaptive coping strategies.')
P(d, 'The programme is designed according to the CBT approach, the stepped care model, and the primary–secondary prevention framework. It can be integrated into morning sessions, flag-raising ceremonies, homeroom activities, thematic sessions, or organised as competitions for lower secondary school students.')

page_break(d)

# CONCLUSION
H1(d, 'CONCLUSION AND RECOMMENDATIONS')

H2(d, '1. Conclusion')
H3(d, '1.1. Theoretical contributions')
P(d, 'The dissertation has systematised the theoretical foundations of anxiety disorders among lower secondary school students through the integration of risk and protective factor networks; established a conceptual framework for anxiety disorders comprising four components (cognitive, emotional, physiological, behavioural); clearly distinguished three forms of anxiety disorder (generalised, social, separation) among lower secondary school students; and established an analytical framework for coping strategies along two dimensions — adaptive and maladaptive.')

H3(d, '1.2. Practical contributions')
P(d, 'The dissertation has provided a detailed picture of the prevalence of anxiety disorders in a sample of 1,352 lower secondary school students in Hanoi. All adapted measurement scales achieved acceptable reliability and structural validity. The level of anxiety disorders among lower secondary school students was primarily at moderate level, with generalised and social anxiety being more prominent than separation anxiety.')
P(d, 'The dissertation has tested and confirmed all three scientific hypotheses:')
P(d, '- Hypothesis 1 (gender differences depend on the type of anxiety disorder): confirmed. Female students showed higher levels of generalised anxiety (M 59.47 vs 51.43; F = 44.484; p < 0.001) and social anxiety (M 52.74 vs 43.20; F = 45.984; p < 0.001) than male students; separation anxiety showed no gender difference (F = 0.246; p = 0.620).', indent=False)
P(d, '- Hypothesis 2 (self-esteem as a strong protective factor): confirmed. Self-esteem showed β = -0.457 for total anxiety disorders, with effect strength approximately 85-89% of academic pressure (per LA Conclusions).', indent=False)
P(d, '- Hypothesis 3 (academic pressure and mobile phone addiction as strong risk factors): confirmed. Order of effect strength: academic pressure (β = 0.533) > mobile phone addiction (β = 0.400) > school bullying (β = 0.276).', indent=False)
P(d, 'Based on these findings, the dissertation proposes a Training Framework for the Prevention of Anxiety Disorders among lower secondary school students with eight core modules, addressing the pressing need for mental health care for Vietnamese students.')

H2(d, '2. Recommendations')
H3(d, '2.1. For students')
P(d, 'Equip students with knowledge and skills for early recognition of anxiety disorder signs; train self-regulation capacity; reduce maladaptive strategies (avoidance, self-blame) through cognitive restructuring (CBT); enhance the quality of adaptive strategies (problem-solving, support-seeking).')

H3(d, '2.2. For teachers and schools')
P(d, 'Integrate the prevention training content into morning sessions, flag-raising, homeroom activities, and thematic sessions. Review and adjust teaching activities that may increase anxiety disorders (assessment pressure, achievement competition). Build teachers\' capacity to recognise anxiety disorder manifestations.')

H3(d, '2.3. For families')
P(d, 'Parents should adjust expectations to align with children\'s capacity and psychological characteristics; build open communication environments; serve as stable and positive sources of psychological support.')

H3(d, '2.4. For school support forces')
P(d, 'School counselling rooms should implement screening and early intervention for at-risk students. Youth unions and clubs should strengthen social connections. Coordination among all forces should follow an integrated, multi-tier intervention model.')

H2(d, '3. Limitations and future research directions')
P(d, 'The study has several limitations: (1) the cross-sectional design cannot confirm causal relationships over time; (2) some measurement structures still require further refinement to enhance stability; (3) self-report methods may be affected by social desirability bias; (4) the sample is limited to two lower secondary schools in Hanoi; (5) the model has not fully examined mediating and moderating variables.')
P(d, 'Future research directions: (1) longitudinal designs across time; (2) further refinement of measurement instruments across different samples; (3) experimental studies to test the effectiveness of the proposed Training Framework for the Prevention of Anxiety Disorders in school settings.')

page_break(d)

# PUBLICATIONS
H1(d, 'LIST OF AUTHOR\'S PUBLICATIONS RELATED TO THE DISSERTATION')

P(d, '[The doctoral candidate is to fill in the list of published articles. Minimum requirement: 2 articles published in scientific journals with citation indices, as required by the Ministry of Education and Training and Hanoi National University of Education.]', italic=True, size=10)
d.add_paragraph()
P(d, '1. Cong Thi Hang. (20...). [Article title]. Journal of ......, Vol. ......, No. ......, pp. ........ DOI: .....................', indent=False, size=11)
d.add_paragraph()
P(d, '2. Cong Thi Hang. (20...). [Article title]. Journal of ......, Vol. ......, No. ......, pp. ........ DOI: .....................', indent=False, size=11)
d.add_paragraph()
P(d, '3. Cong Thi Hang. (20...). [Article title]. Journal of ......, Vol. ......, No. ......, pp. ........ DOI: .....................', indent=False, size=11)

cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''; cp.keywords = ''
cp.comments = ''; cp.last_modified_by = ''

d.save(OUT)
from docx import Document as D
dd = D(OUT)
w = sum(len(p.text.split()) for p in dd.paragraphs)
print(f"Saved: {OUT}")
print(f"Paragraphs: {len(dd.paragraphs)}, Tables: {len(dd.tables)}, Words: ~{w}, Size: {os.path.getsize(OUT)//1024}KB")
