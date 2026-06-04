# -*- coding: utf-8 -*-
"""Tạo 11 tóm tắt theo mẫu 00_Mẫu tóm tắt bài 1.docx"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

def new_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)
    return doc

def red_bold(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.color.rgb = RED

def blue(doc, text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.font.color.rgb = BLUE
    r.bold = bold

def blue_h3(doc, text):
    h = doc.add_heading(text, level=3)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = BLUE

def red_h2(doc, text):
    h = doc.add_heading(text, level=2)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RED

papers = [
    {
        'num': 1, 'file': '01_Jenkins_2023',
        'header': 'Công trình nghiên cứu \u00ab Trầm cảm và lo âu ở học sinh trung học cơ sở đa sắc tộc: Tuổi, giới tính và môi trường văn hóa xã hội \u00bb (Depression and anxiety among multiethnic middle school students: Age, gender, and sociocultural environment), do Janis H. Jenkins và cs. (2023) khảo sát trên mẫu 75 học sinh trung học cơ sở đa sắc tộc từ 10 đến 14 tuổi tại một trường công lập ở San Diego, California, Hoa Kỳ.',
        'method': 'Công trình này sử dụng các công cụ sàng lọc PHQ-9A (Patient Health Questionnaire-9 Adolescent \u2014 Bảng hỏi Sức khỏe Bệnh nhân phiên bản thanh thiếu niên) và GAD-10 (Generalized Anxiety Disorder 10-item Scale \u2014 Thang Rối loạn Lo âu Tổng quát 10 mục) kết hợp với phỏng vấn dân tộc học bán cấu trúc. Nói cách khác, công trình này sử dụng phương pháp hỗn hợp.',
        'tools': [
            'PHQ-9A là phiên bản sửa đổi cho thanh thiếu niên của thang sàng lọc trầm cảm Bảng câu hỏi Sức khỏe Bệnh nhân (Johnson và cộng sự, 2002).',
            'GAD-10 là phiên bản dành cho trẻ em (tuổi 11\u201317) của thang Rối loạn Lo âu Tổng quát (Craske và cộng sự, 2013).',
        ],
        'lit_review': 'Tổng quan tài liệu của chúng tôi cho thấy các thang PHQ-9 và GAD-7 dành cho người lớn thường được sử dụng không phù hợp với trẻ em và thanh thiếu niên sớm (Cortez và cộng sự, 2021; Aggarwal và cộng sự, 2017; Mossman và cộng sự, 2017). Vì vậy, nghiên cứu sử dụng các phiên bản sửa đổi phù hợp lứa tuổi.',
        'demographics': 'Dữ liệu nhân khẩu xã hội bao gồm tuổi, bản dạng giới tính, thành phần gia đình và tôn giáo.',
        'interview': 'Hướng dẫn phỏng vấn sâu của chúng tôi được thiết kế để tìm hiểu một cách mở, với sự linh hoạt cho các câu hỏi thăm dò. Các chủ đề phỏng vấn bao gồm cuộc sống hàng ngày, trải nghiệm gia đình và trường học, sức khỏe tâm thần và sức khỏe tổng thể, quan điểm về bạn bè đồng trang lứa, cũng như mục tiêu tương lai. Tất cả các cuộc phỏng vấn được tiến hành bằng tiếng Anh và tiếng Tây Ban Nha, được ghi âm để phiên âm và phân tích mã hóa chủ đề định tính.',
        'mixed_method': 'Phương pháp hỗn hợp của chúng tôi sử dụng các thang sàng lọc trầm cảm và lo âu kết hợp với các công cụ phỏng vấn dân tộc học, được xây dựng dựa trên các nghiên cứu trước đây chứng minh tính bổ sung của nhiều loại bằng chứng nghiên cứu khác nhau soi chiếu lẫn nhau trên các lĩnh vực tìm hiểu (Csordas và cộng sự, 2010; Jenkins, 2015). Ví dụ, trong khi học sinh có thể phủ nhận sự hiện diện của các triệu chứng khi được đánh giá theo thang điểm, họ có thể kể lại trải nghiệm của mình trong một cuộc phỏng vấn mở. Ngược lại, một số người tham gia nghiên cứu có thể xác nhận các mục trên thang đo mà họ có thể không muốn nói đến trong phỏng vấn tường thuật.',
        'results': 'Về tần suất, mức "nhẹ" là phổ biến nhất : trầm cảm nhẹ là 22,6% và lo âu nhẹ là 32%. Mức trung bình đến cực kỳ nặng ít được quan sát thấy hơn. Khi gộp lại bất kỳ mức độ triệu chứng nào, khoảng một nửa (50,6%) học sinh trong nghiên cứu rơi vào phạm vi từ nhẹ đến nặng của thang sàng lọc lo âu GAD-10, trong khi 44% rơi vào phạm vi từ nhẹ đến nặng của thang sàng lọc trầm cảm PHQ-9A.',
        'gender': 'Khác biệt giới tính : Nữ giới luôn báo cáo mức trầm cảm cao hơn nam giới (Salk và cộng sự, 2017). Thực tế, sự chênh lệch giới trong trầm cảm đã mở rộng ở thanh thiếu niên sớm tại Hoa Kỳ trong thập kỷ qua (Daly, 2022). Chúng tôi xác nhận trong nhóm thanh thiếu niên sớm này rằng điểm PHQ-9A cao hơn ở nữ so với nam (p = 0,002, kiểm định Mann-Whitney), với chênh lệch trung vị là 4 điểm. Tương tự, điểm GAD-10 cao hơn ở nữ so với nam (p = 0,016).',
        'qualitative': [
            ('COVID-19', 'Nổi bật đối với cả nữ sinh và nam sinh là sự lo lắng về hậu quả của COVID-19 (Bệnh virus corona 2019 \u2014 Coronavirus Disease 2019) đối với gia đình liên quan đến bệnh tật, tử vong, mất việc làm và khó khăn kinh tế. Đây là một lĩnh vực sức khỏe tâm thần không nằm trong các mục sàng lọc của thang PHQ-9A và GAD-10 nhưng được mô tả chi tiết bởi các cuộc phỏng vấn dân tộc học.'),
            ('Phân biệt đối xử', 'Nỗi sợ hãi trước nhận thức của học sinh về sự gia tăng phân biệt đối xử và phân biệt chủng tộc trong trường học và cộng đồng là rất nổi bật. Những vấn đề này được học sinh trải nghiệm như những rào cản đối với sự tham gia giáo dục.'),
            ('Bạo lực trên cơ sở giới', 'Các cuộc phỏng vấn dân tộc học cho thấy các bối cảnh liên quan đến trải nghiệm trầm cảm và lo âu của nữ sinh, bao gồm bạo lực trên cơ sở giới trong cả môi trường trường học và gia đình.'),
        ],
        'conclusion': 'Dữ liệu của chúng tôi, cho thấy sự hiện diện của các mức độ triệu chứng trầm cảm và lo âu từ nhẹ đến nặng trong một mẫu phi lâm sàng gồm các học sinh trung học cơ sở đa sắc tộc, gợi ý rằng cần chú ý nghiên cứu ở độ tuổi sớm hơn để nhận diện tình trạng dưới ngưỡng hội chứng. Sự khác biệt giới tính có ý nghĩa, trước đây được tìm thấy ở người trưởng thành và thanh thiếu niên lớn tuổi hơn, đã có thể quan sát được trong nhóm tuổi 10\u201314 của nghiên cứu này. Trong bối cảnh trường trung học cơ sở thu nhập thấp, đa sắc tộc như nghiên cứu hiện tại, nhu cầu giải quyết mối quan hệ giữa các yếu tố căng thẳng về văn hóa xã hội và sức khỏe tâm thần, dù ở mức nhẹ hay nặng, là thiết yếu để tạo điều kiện cho sự tham gia giáo dục.',
        'critique': 'Cỡ mẫu rất nhỏ (n=75), chọn mẫu thuận tiện tại 1 trường duy nhất \u2014 không đại diện. GAD-10 là phiên bản sửa đổi thiếu dữ liệu xác thực (validation). Thiết kế kéo dài 2018\u20132021 trùng COVID-19 nhưng không kiểm soát hiệu ứng thời điểm. Số phỏng vấn dân tộc học không đồng đều (1\u20138/học sinh) gây thiên lệch. Tuy nhiên, phương pháp hỗn hợp là điểm sáng \u2014 ít nghiên cứu nào kết hợp dân tộc học với sàng lọc.',
        'gap': 'Cần mở rộng cỡ mẫu để xác nhận kết quả. Phát triển công cụ sàng lọc nhạy cảm văn hóa cho thanh thiếu niên sớm đa sắc tộc. Nghiên cứu dọc theo dõi sự phát triển triệu chứng từ đầu vị thành niên. Đánh giá hiệu quả can thiệp giải quyết yếu tố văn hóa xã hội đặc thù.',
        'quality': '\u2b50\u2b50\u2b50 Trung bình-Khá. Phương pháp hỗn hợp sáng tạo nhưng cỡ mẫu nhỏ hạn chế khả năng tổng quát hóa.',
    },
]

# Tạo bài 1 theo mẫu
for p in papers:
    doc = new_doc()

    # Header
    blue(doc, f"Tóm tắt bài {p['num']}", bold=True)
    red_bold(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
    blue(doc, p['header'])

    # Phương pháp
    red_bold(doc, 'Phương pháp nghiên cứu')
    blue(doc, p['method'], bold=True)
    for tool in p['tools']:
        blue(doc, tool)
    blue(doc, p['lit_review'], bold=True)
    blue(doc, p['demographics'])
    blue(doc, p['interview'])
    blue(doc, p['mixed_method'])

    # Kết quả
    red_bold(doc, 'Kết quả nghiên cứu định lượng')
    blue(doc, p['results'])
    blue_h3(doc, p['gender'])

    # Nhận xét
    red_bold(doc, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
    for label, text in p['qualitative']:
        blue(doc, f'*{label}.* {text}')

    # Kết luận
    red_bold(doc, 'Kết luận')
    blue(doc, p['conclusion'], bold=True)

    # Phản biện
    red_h2(doc, 'Phản biện')
    blue(doc, p['critique'])

    # Hướng NC
    red_h2(doc, 'Hướng nghiên cứu tiếp theo')
    blue(doc, p['gap'])

    # Chất lượng
    doc.add_paragraph()
    pp = doc.add_paragraph()
    r = pp.add_run(f"Đánh giá chất lượng: {p['quality']}")
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)

    fname = f"{p['file']}.docx"
    doc.save(fname)
    sys.stderr.write(f'{fname} OK\n')

sys.stderr.write('Done bài 1\n')
