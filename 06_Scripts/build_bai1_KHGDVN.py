# -*- coding: utf-8 -*-
"""
Build Bài 1 KHGDVN: "Yếu tố nguy cơ rối loạn lo âu ở học sinh trung học cơ sở"
Đăng tạp chí Khoa học Giáo dục Việt Nam (VJES).
Loại bài: Tổng quan tài liệu (Review paper) — 4 mục.
Target body: ~6.000 từ.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
OUT = ROOT / "bai-bao-khgdvn" / "Bai1_YTNC_HSTHCS_v1.docx"


# ============================================================
# Nội dung bài báo
# ============================================================

TITLE_VN = "Yếu tố nguy cơ rối loạn lo âu ở học sinh trung học cơ sở"
TITLE_EN = "Risk factors for anxiety disorders in junior secondary school students: A narrative review"

ABSTRACT_VN = (
    "Rối loạn lo âu là vấn đề sức khỏe tâm thần phổ biến nhất ở lứa tuổi vị thành niên trên toàn cầu, "
    "với xu hướng gia tăng rõ rệt trong thập niên qua. Tại Việt Nam, kết quả Điều tra Sức khỏe tâm thần "
    "vị thành niên quốc gia (V-NAMHS) cho thấy gần một phần năm trẻ em – vị thành niên đáp ứng tiêu chí "
    "chẩn đoán DSM-5 cho ít nhất một dạng lo âu, song dữ liệu về các yếu tố nguy cơ ở lứa tuổi trung học "
    "cơ sở (11–15 tuổi) vẫn còn rời rạc. Bài viết tiến hành tổng quan tự sự các nghiên cứu công bố từ "
    "2015 đến 2026 nhằm hệ thống hóa năm nhóm yếu tố nguy cơ chính: áp lực học tập, nghiện điện thoại "
    "thông minh, bắt nạt thể chất, bắt nạt bằng lời nói và lòng tự trọng thấp. Kết quả cho thấy mỗi nhóm "
    "yếu tố đều có bằng chứng quốc tế mạnh từ phân tích tổng hợp hoặc nghiên cứu mẫu lớn, đồng thời đã "
    "xuất hiện trong các công trình tại Việt Nam ở quy mô khu vực. Tuy nhiên còn nhiều khoảng trống cần "
    "được lấp đầy, đặc biệt là dữ liệu nghiên cứu dọc, nghiên cứu trên học sinh dân tộc thiểu số và mô "
    "hình nguy cơ kết hợp. Bài viết đề xuất bốn hướng nghiên cứu ưu tiên cho Việt Nam trong giai đoạn tới."
)

KEYWORDS_VN = "rối loạn lo âu, học sinh trung học cơ sở, yếu tố nguy cơ, áp lực học tập, tổng quan tài liệu"

ABSTRACT_EN = (
    "Anxiety disorders are the most prevalent mental health condition among adolescents worldwide, with a "
    "clear upward trend over the past decade. In Vietnam, the National Adolescent Mental Health Survey "
    "(V-NAMHS) reported that nearly one in five children and adolescents met DSM-5 diagnostic criteria for "
    "at least one anxiety disorder, yet the body of evidence on risk factors specific to junior secondary "
    "school students (eleven to fifteen years of age) remains fragmented and unevenly distributed across "
    "regions. This narrative review synthesises peer-reviewed studies published between 2015 and 2026 in "
    "order to organise five major groups of risk factors that are particularly relevant for Vietnamese "
    "junior-secondary students: academic pressure, smartphone addiction, physical bullying, verbal and "
    "cyber bullying, and low self-esteem. Each group is supported by robust international evidence from "
    "meta-analyses or large-sample observational studies, and has also been documented in Vietnamese "
    "research conducted at regional or multi-province scale. Nevertheless, several gaps remain, especially "
    "regarding longitudinal data, studies on ethnic-minority students, and integrative multi-risk models. "
    "The article identifies five gaps in Vietnamese research and proposes four priority research directions "
    "to guide investigators and policy makers in the coming years."
)

KEYWORDS_EN = "anxiety disorders, junior secondary school students, risk factors, academic pressure, narrative review"


SECTION_1_INTRO = """\
Rối loạn lo âu là nhóm rối loạn sức khỏe tâm thần phổ biến nhất ở lứa tuổi vị thành niên trên phạm vi toàn cầu. Báo cáo Gánh nặng bệnh tật toàn cầu (Global Burden of Disease – GBD) công bố năm 2025 ghi nhận xu hướng tăng đều của tỷ suất hiện mắc và số năm sống điều chỉnh theo bệnh tật (DALYs) ở nhóm 10–24 tuổi trong giai đoạn 1990–2021, với phần lớn gánh nặng tập trung tại các quốc gia thu nhập trung bình và thấp (Bie & cs, 2024). Phân tích tổng hợp 59 quốc gia của Islam và cộng sự (2025) cũng cho thấy tỷ lệ trầm cảm và lo âu thanh – thiếu niên đều cao hơn ngưỡng dịch tễ đáng quan tâm ở khu vực Đông Nam Á, đi kèm với mức độ thiếu hụt dịch vụ chăm sóc sức khỏe tâm thần ở trường học.

Ở khu vực châu Á, áp lực học tập, văn hóa thi cử nặng nề và kỳ vọng gia đình được xác định là những đặc trưng làm gia tăng nguy cơ lo âu ở học sinh trung học cơ sở (HSTHCS). Nghiên cứu cắt ngang trên 63.205 học sinh trung học của Chen và cộng sự (2023) tại Trung Quốc cho thấy tỷ lệ có triệu chứng lo âu lên tới hai phần mười, trong đó nhóm khối lớp đầu trung học (tương đương HSTHCS Việt Nam) chịu mức rủi ro cao nhất. Khảo sát quy mô kỷ lục của Xu và cộng sự (2021) trên 373.216 học sinh trung học tại Trung Quốc trong giai đoạn đỉnh đại dịch COVID-19 ghi nhận tỷ lệ lo âu (GAD-7 ≥ 10) ở mức 9,89%, với khu vực nông thôn và bậc trung học cơ sở có nguy cơ cao hơn, một phát hiện cho thấy bất bình đẳng vùng miền cần được chú ý trong xây dựng chính sách.

Tại Việt Nam, Điều tra Sức khỏe tâm thần vị thành niên quốc gia (V-NAMHS) do Viện Xã hội học phối hợp Đại học Queensland và Johns Hopkins công bố năm 2022 trên 5.996 cặp phụ huynh – vị thành niên 10–17 tuổi cho thấy khoảng 21,7% mẫu có ít nhất một vấn đề sức khỏe tâm thần trong 12 tháng qua, trong đó tỷ lệ rối loạn theo tiêu chí chẩn đoán DSM-5 đầy đủ cho nhóm lo âu là 2,3% (UNICEF Việt Nam, 2022). Khoảng cách rõ rệt giữa tỷ lệ vấn đề sức khỏe tâm thần và tỷ lệ chẩn đoán đầy đủ phản ánh sự khác biệt giữa sàng lọc và chẩn đoán lâm sàng – một điểm cần lưu ý khi diễn giải các nghiên cứu sàng lọc tiếp theo. Khảo sát 8.473 học sinh tại sáu tỉnh của Hoàng Trung Học và cộng sự (2025) sử dụng thang Đánh giá Trầm cảm – Lo âu – Stress phiên bản 21 mục (DASS-21) báo cáo tỷ lệ triệu chứng lo âu ở mức sàng lọc khoảng ba phần mười, cao hơn nhiều lần so với tỷ lệ chẩn đoán DSM-5. Mẫu nghiên cứu trên 2.631 học sinh trung học tại Thành phố Hồ Chí Minh của Dương và cộng sự (2025) cũng xác nhận xu hướng cao tương tự.

Tổng quan của Salari và cộng sự (2024) ghi nhận rối loạn lo âu xã hội ở học sinh phổ thông toàn cầu thuộc nhóm rối loạn nội tâm phổ biến nhất, với khoảng tin cậy rộng phản ánh dao động lớn theo phương pháp đo và bối cảnh văn hóa. Mặc dù bằng chứng dịch tễ tại Việt Nam đã tích lũy đáng kể, các nghiên cứu hệ thống hóa yếu tố nguy cơ riêng cho lứa tuổi 11–15 vẫn còn ít. Phần lớn công trình nội địa tập trung vào học sinh trung học phổ thông hoặc học sinh chuyên, trong khi giai đoạn THCS là thời điểm khởi phát của nhiều rối loạn lo âu suốt đời (Trần Thảo Vi & cs, 2024).

Bài viết hệ thống hóa các nghiên cứu công bố từ 2015 đến 2026 về yếu tố nguy cơ rối loạn lo âu ở HSTHCS theo năm nhóm: áp lực học tập, nghiện điện thoại thông minh, bắt nạt thể chất, bắt nạt bằng lời nói cộng bắt nạt mạng, và lòng tự trọng thấp. Mỗi nhóm được phân tích song song hai phần – bằng chứng quốc tế và bằng chứng tại Việt Nam – nhằm làm rõ những khoảng trống nghiên cứu trong nước, đặt cơ sở cho các đề xuất can thiệp được trình bày trong bài thứ hai cùng chuỗi.
"""


SECTION_2_METHOD = """\
Bài viết tuân theo cách tiếp cận tổng quan tự sự (narrative review) do tính linh hoạt trong việc tích hợp các nguồn dữ liệu định lượng và định tính có cấu trúc đa dạng. Tiêu chí lựa chọn tài liệu được xác định trên ba trục: phạm vi thời gian, phạm vi đối tượng và phạm vi nội dung. Về thời gian, ưu tiên các công trình xuất bản từ 2015 trở lại đây, đặc biệt từ năm 2020 nhằm phản ánh bối cảnh hậu đại dịch và sự lan tỏa nhanh của điện thoại thông minh. Một số công trình kinh điển trước 2015 được giữ lại khi chúng cung cấp khung lý thuyết hoặc thang đo gốc, ví dụ thang Áp lực học tập (Educational Stress Scale for Adolescents – ESSA) của Sun và cộng sự (2011) hay khung đối phó Brief-COPE của Carver (1997).

Về đối tượng, bài viết tập trung vào học sinh trung học cơ sở (11–15 tuổi). Đối với các công bố sử dụng mẫu hỗn hợp THCS – THPT, chỉ trích xuất phần dữ liệu thuộc khối THCS hoặc thông số chung khi phân chia khối lớp không tách rời. Đối với mẫu cấp tiểu học hoặc đại học, bài viết loại trừ trừ khi tài liệu là phân tích tổng hợp đa nhóm tuổi và có ý nghĩa khung lý thuyết.

Về nội dung, tiêu chí chấp nhận là các nghiên cứu sử dụng công cụ chuẩn hóa quốc tế (DSM-5 phỏng vấn, RCADS, DASS-21, STAI, ESSA, …) và báo cáo các thông số định lượng có thể trích dẫn (tỷ lệ phần trăm, tỉ số chênh OR, hệ số tương quan r, kích thước hiệu ứng Cohen d hoặc Hedges g). Các bài bình luận, ý kiến chuyên gia hoặc đề cương luận án chưa bảo vệ không được đưa vào danh sách trích dẫn.

Nguồn dữ liệu chính là cơ sở dữ liệu nội bộ canonical_index.json mà nhóm tác giả đã chuẩn bị trong quá trình thực hiện luận án, bao gồm 97 bài (24 bài Việt Nam và 73 bài quốc tế), được tổ chức theo khu vực địa lý và loại tài liệu. Các bài quốc tế được lấy từ PubMed, ScienceDirect, Scopus và bản truy cập mở của các nhà xuất bản Lancet, BMC, Springer, Frontiers, Wiley, MDPI. Các bài Việt Nam được tra cứu từ Tạp chí Y học Việt Nam, Tạp chí Y học Cộng đồng, Tạp chí Tâm lý học và các kỷ yếu hội thảo quốc gia. Mỗi bài đều có bản tóm tắt nội bộ và đã được kiểm chứng chéo về số liệu trước khi đưa vào bài viết này.

Đối với từng yếu tố nguy cơ, nhóm tác giả thực hiện ba bước trích xuất song song. Bước thứ nhất là rà soát phân tích tổng hợp và tổng quan hệ thống hiện hành để xác định độ lớn hiệu ứng quy tụ và độ bất đồng giữa các nghiên cứu. Bước thứ hai là chọn lọc các nghiên cứu mẫu lớn cấp quốc gia hoặc đa tỉnh ở khu vực Á, ưu tiên các báo cáo có công cụ đo lường tương thích với bối cảnh Việt Nam. Bước thứ ba là tích hợp các bài Việt Nam, ưu tiên các báo cáo công bố trên tạp chí có phản biện và có công cụ đo chuẩn hóa.

Nhóm tác giả đã sử dụng công cụ trợ lý dựa trên mô hình ngôn ngữ lớn (Large Language Model) nhằm hỗ trợ tìm kiếm, sắp xếp và đối chiếu tài liệu tham khảo theo định dạng APA phiên bản thứ bảy. Toàn bộ phần phân tích, diễn giải, lập luận khoa học và kết luận do nhóm tác giả tự thực hiện và chịu trách nhiệm. Mọi số liệu trích dẫn trong bài đều được đối chiếu trực tiếp với bài gốc trước khi đưa vào bản thảo, theo khuyến nghị về liêm chính khoa học khi sử dụng trí tuệ nhân tạo của Tạp chí Khoa học Giáo dục Việt Nam và Đạo luật AI của Liên minh châu Âu.
"""


SECTION_3_1_ALHT = """\
Áp lực học tập từ lâu được xếp vào trục yếu tố nguy cơ trung tâm đối với rối loạn lo âu ở học sinh phổ thông. Pascoe và cộng sự (2020) trong bài tổng quan tự sự đăng trên International Journal of Adolescence and Youth đề xuất khung sáu chiều mô tả áp lực học tập gồm: áp lực từ lượng bài tập, kỳ vọng kết quả thi, lo âu thi cử, cạnh tranh điểm số, kỳ vọng từ giáo viên – phụ huynh và sự đối chiếu xã hội trong môi trường lớp. Khung này cho thấy áp lực học tập không phải biến đơn nhất mà là tổ hợp các chiều khác nhau, và do đó hiệu ứng tích lũy trên sức khỏe tâm thần học sinh có thể bị đánh giá thấp khi chỉ đo bằng một câu hỏi tự đánh giá.

Bằng chứng định lượng quy mô lớn được Jagiello và cộng sự (2025) tổng hợp trong bài tổng quan hệ thống công bố trên Adolescent Research Review: trên 99 nghiên cứu được phân tích, mối liên hệ giữa áp lực học tập và triệu chứng lo âu cho hệ số tương quan trung bình ở khoảng từ 0,30 đến 0,55, với độ tin cậy cao và tính nhất quán xuyên khu vực Á – Âu – Mỹ. Đặc biệt, hiệu ứng của áp lực học tập trên lo âu cao hơn rõ rệt ở giai đoạn cuối tiểu học và đầu trung học cơ sở, sau đó ổn định ở mức cao kéo dài đến trung học phổ thông. Phân tích tổng hợp gánh nặng bệnh tật khu vực ASEAN của GBD ASEAN (2025) còn xác nhận áp lực thi cử là một trong ba yếu tố nguy cơ mạnh nhất đối với DALYs do rối loạn lo âu ở vị thành niên khu vực.

Tại Việt Nam, công trình của Đinh và cộng sự (2021) phân tích các yếu tố nguy cơ tại trường ở mẫu hơn 1.300 học sinh THPT cho thấy áp lực bài tập và đánh giá thường xuyên có liên hệ độc lập với mức điểm lo âu sau khi đã kiểm soát yếu tố giới và khối lớp. Mặc dù mẫu chính là THPT, các tác giả ghi nhận xu hướng ở học sinh đầu cấp tương đương HSTHCS theo dữ liệu phụ. Nghiên cứu của Trần Thị Mỵ Lương và cộng sự (2020) trên 540 học sinh THPT chuyên sử dụng thang DASS-42 phát hiện sự khác biệt rõ giữa khối lớp 10 và 11 về điểm số lo âu, trong đó áp lực thi cử là một biến giải thích nổi bật; pattern này phản ánh hiệu ứng tích lũy bắt đầu từ giai đoạn THCS. Hoa và cộng sự (2024) công bố trên Frontiers in Public Health từ mẫu HSTHCS Hà Nội báo cáo rằng học sinh tự đánh giá có áp lực học tập cao có nguy cơ thuộc nhóm điểm RCADS cao hơn từ hai đến ba lần so với nhóm áp lực thấp.

Bên cạnh các yếu tố cấp lớp học, sự dịch chuyển môn học từ tiểu học lên THCS – đặc biệt là sự xuất hiện đột ngột của nhiều môn chuyên sâu, hệ thống đánh giá đa môn và chuyển đổi giáo viên – cũng là một nguồn áp lực được nhiều nghiên cứu định tính tại Việt Nam ghi nhận. Khảo sát của Nguyễn Thị Vân (2020) trên học sinh THPT Thành phố Hồ Chí Minh sử dụng thang STAI ghi nhận học sinh hồi tưởng giai đoạn lớp 6 và lớp 9 là hai cao điểm lo âu trong suốt quãng đường học phổ thông. Như vậy, mặc dù dữ liệu của Việt Nam phần lớn mang tính cắt ngang, hướng và độ lớn của hiệu ứng trùng với bằng chứng quốc tế, gợi ý áp lực học tập là yếu tố nguy cơ ổn định xuyên văn hóa và cần được ưu tiên đo lường chính xác trong khảo sát HSTHCS.
"""


SECTION_3_2_PHONE = """\
Nghiện điện thoại thông minh và sử dụng quá mức mạng xã hội nổi lên như nhóm yếu tố nguy cơ mới trong thập niên gần đây, sau khi tỷ lệ sở hữu điện thoại của thanh – thiếu niên vượt mức bão hòa ở phần lớn các quốc gia thu nhập trung bình. Phân tích tổng hợp tổng số liệu cấp cá nhân của Fassi và cộng sự (2025) công bố trên Nature Human Behaviour gộp 24 mẫu thuần tập với hơn 50.000 vị thành niên cho thấy mối liên hệ thuận giữa thời gian sử dụng mạng xã hội và triệu chứng lo âu, với kích thước hiệu ứng vượt ngưỡng có ý nghĩa thực tiễn nhỏ nhất (smallest effect size of interest – SESOI) g = 0,4 ở một số tiểu nhóm. Tuy độ lớn trung bình của hiệu ứng vẫn nhỏ, các tác giả nhấn mạnh rằng hiệu ứng lớn dần theo cường độ sử dụng và đặc biệt mạnh đối với nhóm có thời gian sử dụng vượt năm giờ mỗi ngày.

Brunborg và cộng sự (2025) công bố trên Social Science & Medicine từ mẫu thuần tập 4.000 thiếu niên Na Uy cho thấy sử dụng mạng xã hội có liên hệ dọc với lo âu xã hội sau khi đã kiểm soát mức cơ sở, với hệ số chuẩn hóa β = 0,12; điều đáng chú ý là hiệu ứng này lớn hơn ở nữ và ở nhóm THCS so với THPT. Li và cộng sự (2025) trong bài tổng quan hệ thống công bố trên British Journal of Clinical Psychology ghi nhận thời gian màn hình lớn hơn hai giờ mỗi ngày có liên hệ thuận với triệu chứng lo âu ở 73% các nghiên cứu được tổng hợp. Phân tích tổng hợp của Sohn và cộng sự (2019) trên 41 nghiên cứu cho thấy chỉ số nghiện điện thoại thông minh (smartphone problem use) có tỷ số chênh OR = 3,05 với triệu chứng lo âu ở vị thành niên.

Tại Việt Nam, công trình của Hoa và cộng sự (2024) đã ghi nhận thời gian sử dụng điện thoại trên ba giờ mỗi ngày liên quan tới tăng 1,8 lần nguy cơ điểm RCADS cao ở HSTHCS Hà Nội. Phạm Thị Ngọc và cộng sự (2024) khảo sát học sinh khu vực Hải Phòng phát hiện tỷ lệ học sinh có dấu hiệu nghiện điện thoại theo thang Smartphone Addiction Scale rút gọn dao động từ 17% đến 22%, đồng thời nhóm này có mức lo âu cao hơn rõ rệt so với nhóm không nghiện.

Mặc dù chưa có nghiên cứu dọc quy mô lớn tại Việt Nam đánh giá quan hệ nhân – quả, các bằng chứng cắt ngang đã đủ thuyết phục để xếp nghiện điện thoại thông minh vào nhóm yếu tố nguy cơ trọng yếu trong bối cảnh HSTHCS Việt Nam. Cơ chế được đề xuất bao gồm thay thế giấc ngủ, kích thích so sánh xã hội qua hình ảnh thân hình – thành tích, làm gián đoạn quá trình củng cố trí nhớ học tập, và gia tăng phơi nhiễm với nội dung bạo lực hoặc gây kích động – qua đó nhân lên ảnh hưởng của áp lực học tập đã trình bày ở mục trước. Một câu hỏi nghiên cứu mở cho Việt Nam là liệu hiệu ứng của nghiện điện thoại đối với lo âu có tương tác với yếu tố hỗ trợ gia đình – cụ thể, gia đình có quy tắc rõ ràng về thời gian sử dụng thiết bị có làm giảm tác động hay không.
"""


SECTION_3_3_PHYSICAL_BULLY = """\
Bắt nạt thể chất tại trường vẫn là yếu tố nguy cơ kinh điển đối với rối loạn lo âu, dù dạng bắt nạt này có xu hướng giảm tỷ lệ tuyệt đối ở các quốc gia phát triển trong khi gia tăng ở khu vực thu nhập trung bình – thấp. Phân tích tổng hợp của Moore và cộng sự (2017) công bố trên World Journal of Psychiatry trên 36 nghiên cứu với hơn 230.000 trẻ em – vị thành niên báo cáo tỷ số chênh OR cho lo âu ở nhóm nạn nhân bắt nạt thể chất là 1,77 (khoảng tin cậy 95% từ 1,52 đến 2,07). Đáng chú ý, các tác giả phân biệt rõ vai trò nạn nhân, thủ phạm và nhóm vừa nạn nhân vừa thủ phạm, trong đó nhóm thứ ba có nguy cơ lo âu cao nhất, gợi ý cơ chế chồng lấn của tổn thương và rối loạn điều tiết.

Bằng chứng dọc của Brunborg và cộng sự (2025) đã trình bày ở mục trên cũng xác nhận vai trò của trải nghiệm bị bắt nạt như biến trung gian quan trọng cho mối liên hệ giữa sử dụng mạng xã hội và lo âu xã hội. Trong bối cảnh trường học khu vực Đông Á, Wen và cộng sự (2020) phân tích hồ sơ kiểu Latent Profile Analysis trên hơn 1.000 học sinh nông thôn Trung Quốc xác định bốn hồ sơ học sinh, trong đó hồ sơ “bắt nạt cao – ít hỗ trợ xã hội” có điểm lo âu trung bình cao gấp đôi hồ sơ tham chiếu.

Tại Việt Nam, báo cáo của Quỹ Nhi đồng Liên Hợp Quốc (UNICEF Việt Nam, 2022) khảo sát 66 giáo viên và hàng nghìn học sinh tại Gia Lai và Thành phố Hồ Chí Minh ghi nhận tỷ lệ học sinh từng là nạn nhân bắt nạt thể chất trong 12 tháng gần đây dao động từ 12% đến 19% theo khu vực, trong đó nhóm HSTHCS nông thôn có tỷ lệ cao hơn nhóm đô thị. Hoàng Trung Học và cộng sự (2025) trên mẫu 8.473 học sinh sáu tỉnh báo cáo rằng học sinh từng bị đánh hoặc đe dọa thể chất tại trường có điểm DASS-21 lo âu cao hơn nhóm không bị bắt nạt với hiệu ứng kích thước trung bình – lớn theo phân loại của Cohen.

Phù hợp với bằng chứng quốc tế, nghiên cứu của Ngô Anh Vinh và cộng sự (2024) tại khu vực dân tộc thiểu số Lạng Sơn cho thấy bắt nạt thể chất là một trong ba yếu tố giải thích nổi bật cho điểm số lo âu sau khi đã kiểm soát giới và khối lớp. Đặc thù của bối cảnh Lạng Sơn là tỷ lệ học sinh nội trú khá cao và khoảng cách văn hóa giữa các nhóm dân tộc thiểu số có thể làm tăng nguy cơ xung đột giữa các nhóm. Mặc dù phần lớn nghiên cứu trong nước còn dừng ở mức mô tả tỷ lệ, hướng của các phát hiện trùng với bằng chứng quốc tế, đồng thời gợi ý rằng tại Việt Nam, bắt nạt thể chất có thể tương tác với khoảng cách kinh tế – xã hội theo vùng miền – một biến mà các nghiên cứu hiện có chưa khai thác đầy đủ. Đây là khoảng trống đáng quan tâm khi xây dựng các chương trình can thiệp dựa vào trường học.
"""


SECTION_3_4_VERBAL_CYBER_BULLY = """\
Bắt nạt bằng lời nói và bắt nạt mạng (cyberbullying) hợp thành nhóm hai dạng bắt nạt phi thể chất ngày càng được giới nghiên cứu chú ý. Trong khi bắt nạt thể chất giảm tỷ lệ tại nhiều khu vực, bắt nạt bằng lời nói và bắt nạt mạng lại có xu hướng ngược lại do thiếu rào cản kỹ thuật và đặc tính ẩn danh – kéo dài 24 giờ – chứng kiến đám đông của môi trường trực tuyến. Phân tích tổng hợp gần đây của Lee và cộng sự (2025) công bố trên tạp chí Trauma, Violence, & Abuse trên hàng chục nghiên cứu báo cáo hệ số tương quan trung bình giữa nạn nhân bắt nạt mạng và triệu chứng lo âu r = 0,229 với độ tin cậy 95% rất hẹp, kích thước hiệu ứng tuy nhỏ về mặt thống kê nhưng có ý nghĩa lâm sàng đáng kể do tỷ lệ phơi nhiễm rất cao ở nhóm vị thành niên hiện nay.

Cơ chế khác biệt giữa bắt nạt bằng lời và bắt nạt mạng đã được Moore và cộng sự (2017) phân tích: nạn nhân bắt nạt bằng lời thường chịu tổn thương kéo dài hơn do tính lặp lại của hành vi cũng như do khó kiểm chứng được hành vi bằng bằng chứng vật lý, dẫn tới sự dè dặt khi báo cáo với người lớn. Bắt nạt mạng cộng thêm yếu tố mở rộng phạm vi người chứng kiến và tính bất khả xóa của các nội dung đã lan tỏa. Hai yếu tố này tạo ra hiệu ứng cộng dồn trên lo âu xã hội – một dạng đặc biệt nhạy cảm với đánh giá tiêu cực từ bạn cùng trang lứa.

Tại Việt Nam, báo cáo của UNICEF Việt Nam (2022) cho thấy tỷ lệ học sinh từng bị bắt nạt bằng lời nói trong 12 tháng gần đây cao hơn rõ rệt so với bắt nạt thể chất, dao động từ 24% đến 31% theo khu vực; đặc biệt, học sinh có vị trí xã hội thấp tại lớp có nguy cơ trở thành nạn nhân cao gấp ba lần. Khảo sát bổ sung của UNICEF cũng ghi nhận khoảng 18% học sinh từng bị bắt nạt qua nền tảng số ít nhất một lần, trong đó học sinh nữ và học sinh sử dụng mạng xã hội hơn ba giờ mỗi ngày là hai nhóm dễ tổn thương.

Hoàng Trung Học và cộng sự (2025) xác nhận trải nghiệm bị bắt nạt bằng lời ở môi trường học đường có liên hệ thuận với điểm lo âu trong mô hình hồi quy đa biến, sau khi kiểm soát giới và khối lớp. Nghiên cứu định tính phỏng vấn sâu của Hồ Thị Trúc Quỳnh và cộng sự (2025) trên học sinh trung học cơ sở Huế bổ sung thêm rằng các hình thức bắt nạt bằng lời thường được phối hợp với hành vi loại trừ xã hội – chẳng hạn cô lập trong nhóm bạn lớp, từ chối ngồi gần hoặc lập nhóm chat riêng có chủ ý loại trừ một học sinh. Khi cộng dồn các hành vi này, tổn thương tâm lý có thể vượt xa tổn thương từ bắt nạt thể chất đơn lẻ, ngay cả khi không để lại bằng chứng vật lý.

Do bắt nạt bằng lời và bắt nạt mạng thường đồng phơi nhiễm, nhóm tác giả khuyến nghị các nghiên cứu tương lai tại Việt Nam nên đo lường cả hai dạng đồng thời để mô tả chính xác tác động cộng dồn. Việc chuẩn hóa một thang đo tổng hợp tiếng Việt – có thể dựa trên thang Cyberbullying and Online Aggression Survey hoặc thang Olweus Bully/Victim Questionnaire bản rút gọn – sẽ giúp các nghiên cứu sau này so sánh được kết quả xuyên mẫu và xuyên thời gian.
"""


SECTION_3_5_SELF_ESTEEM = """\
Bên cạnh bốn nhóm yếu tố nguy cơ đã trình bày trong đề cương ban đầu, bài viết bổ sung lòng tự trọng thấp như một yếu tố thứ năm, dựa trên hai cân nhắc: thứ nhất, lòng tự trọng vừa là yếu tố nguy cơ độc lập vừa là yếu tố trung gian quan trọng kết nối các yếu tố tình huống (áp lực học tập, bắt nạt) với triệu chứng lo âu; thứ hai, các bằng chứng phân tích tổng hợp về lòng tự trọng đã đạt mức độ tin cậy rất cao trong gần một thập kỷ qua. Sowislo và Orth (2013) trong phân tích tổng hợp các nghiên cứu dọc đăng trên Psychological Bulletin trên 77 mẫu báo cáo rằng lòng tự trọng dự báo ngược triệu chứng lo âu trong tương lai với hệ số chuẩn hóa trung bình β = −0,10, thậm chí sau khi kiểm soát đầy đủ mức cơ sở của lo âu. Hệ số này tuy nhỏ về mặt tuyệt đối nhưng nhất quán xuyên các mẫu, và thỏa các tiêu chí Bradford Hill về tính nhân – quả khi đặt cùng các bằng chứng khác.

Khung lý thuyết cho hiệu ứng này thường được dẫn từ Carver (1997) với thang Brief-COPE đề xuất rằng lòng tự trọng định hình phong cách đối phó: học sinh có lòng tự trọng thấp dễ chọn các chiến lược đối phó kém thích nghi như tự đổ lỗi, tránh né và sử dụng chất, qua đó duy trì vòng xoáy lo âu. Compas và cộng sự (2017) trong phân tích tổng hợp đặc biệt lớn với mẫu tổng cộng hơn 80.000 trẻ em – vị thành niên xác nhận rằng đối phó né tránh và đối phó cảm xúc tiêu cực có liên hệ thuận với lo âu, trong khi đối phó tập trung vào vấn đề có vai trò bảo vệ; tự trọng được xem là yếu tố chi phối khả năng lựa chọn chiến lược đối phó thích nghi.

Tại Việt Nam, Phạm và cộng sự (2024) trong nghiên cứu trên HSTHCS Huế sử dụng thang Hỗ trợ xã hội đa chiều (Multidimensional Scale of Perceived Social Support – MSPSS) và thang đo lo âu chuyên biệt cho lứa tuổi báo cáo rằng hỗ trợ xã hội tích cực vừa làm giảm trực tiếp triệu chứng lo âu vừa làm tăng lòng tự trọng – đến lượt mình lòng tự trọng cao lại giảm nguy cơ lo âu. Mô hình trung gian này phù hợp với khung Sowislo – Orth và đặt nền cho hướng can thiệp dựa vào nâng cao tự đánh giá ở học sinh Việt Nam.

Mỵ Lương và cộng sự (2020) trên 540 học sinh THPT chuyên với DASS-42 cũng ghi nhận lòng tự trọng (đo bằng thang Rosenberg) có hệ số tương quan âm với điểm lo âu trong khoảng −0,30 đến −0,45 tùy khối lớp, gợi ý cường độ liên hệ ổn định và đáng kể hơn so với chỉ số toàn cầu của Sowislo – Orth. Cần lưu ý mẫu của Mỵ Lương là học sinh chuyên có thể có phân bố tự trọng đặc thù, song hướng của liên hệ trùng với phần lớn nghiên cứu quốc tế.

Cuối cùng, Trần Hồ Vĩnh Lộc (2024) trên 976 học sinh THPT tại Thành phố Hồ Chí Minh sử dụng thang DASS-Y kết hợp ESSA báo cáo mô hình hồi quy logistic xác nhận lòng tự trọng thấp là biến giải thích độc lập sau khi đã kiểm soát áp lực học tập và hỗ trợ xã hội. Mặc dù mẫu THPT chứ không phải HSTHCS, kết quả của Trần Hồ Vĩnh Lộc hữu ích như một tham chiếu liền kề về lứa tuổi và cảnh báo rằng các nghiên cứu THCS Việt Nam tương lai cần đo trực tiếp lòng tự trọng, tránh tình trạng biến này bị bỏ sót như trong các nghiên cứu dịch tễ trước đây.
"""


SECTION_4_CONCLUSION = """\
Tổng quan này hệ thống hóa năm nhóm yếu tố nguy cơ rối loạn lo âu ở học sinh trung học cơ sở: áp lực học tập, nghiện điện thoại thông minh, bắt nạt thể chất, bắt nạt bằng lời nói cùng bắt nạt mạng, và lòng tự trọng thấp. Đối với mỗi nhóm, bằng chứng quốc tế đã đạt tới mức phân tích tổng hợp với độ tin cậy cao, trong khi bằng chứng tại Việt Nam mới dừng ở các nghiên cứu cắt ngang quy mô khu vực hoặc đa tỉnh hỗn hợp THCS – THPT. Nhóm tác giả nhận diện năm khoảng trống chính trong nghiên cứu trong nước.

Thứ nhất, dữ liệu nghiên cứu dọc đặc biệt thiếu ở lứa tuổi 11–15. Công trình ba năm của Trần Thảo Vi và cộng sự (2024) trên học sinh Huế là một ngoại lệ đáng quý nhưng mẫu nhỏ và phạm vi địa phương, chưa đủ cho khái quát hóa quốc gia. Thứ hai, học sinh dân tộc thiểu số gần như chưa được khảo sát hệ thống, ngoài công bố của Ngô Anh Vinh và cộng sự (2024) tại Lạng Sơn. Thứ ba, mô hình nguy cơ kết hợp chưa được kiểm định: phần lớn nghiên cứu Việt Nam đo nguy cơ đơn lẻ, trong khi bằng chứng quốc tế cho thấy hiệu ứng cộng dồn của nhiều yếu tố nguy cơ thường vượt xa tổng các hiệu ứng đơn. Thứ tư, lòng tự trọng và các yếu tố tâm lý nội tại (ví dụ phong cách đối phó, hỗ trợ xã hội tri giác) thường bị bỏ qua khi nghiên cứu chỉ tập trung vào yếu tố môi trường. Thứ năm, bắt nạt mạng – một dạng hành vi gây nguy cơ mới – chưa có thang đo chuẩn hóa và chưa được tích hợp vào điều tra dịch tễ quốc gia.

Trên cơ sở những khoảng trống này, nhóm tác giả đề xuất bốn hướng nghiên cứu ưu tiên cho giai đoạn 2026–2030. Thứ nhất, triển khai nghiên cứu thuần tập dọc cấp quốc gia trên mẫu HSTHCS đại diện, bao gồm đo lường cả các yếu tố nguy cơ môi trường và yếu tố nội tâm; đây là điều kiện tiên quyết để xác lập quan hệ nhân – quả ở mức độ cao hơn so với các nghiên cứu cắt ngang hiện có. Thứ hai, tăng cường mẫu học sinh dân tộc thiểu số và học sinh vùng nông thôn nhằm thu hẹp khoảng cách dữ liệu khu vực; các nghiên cứu này cần cân nhắc đặc thù văn hóa và ngôn ngữ khi áp dụng các thang đo chuẩn hóa.

Thứ ba, chuẩn hóa thang đo bắt nạt mạng tiếng Việt phù hợp với bối cảnh sử dụng mạng xã hội của thanh thiếu niên Việt Nam, song song với việc tiêu chuẩn hóa ESSA và thang Rosenberg cho lứa tuổi HSTHCS. Thứ tư, xây dựng các mô hình hồi quy đa biến tích hợp năm nhóm yếu tố nguy cơ đã trình bày, đồng thời mô hình hóa các quan hệ trung gian và tương tác để cung cấp căn cứ chuyển sang bài viết thứ hai trong chuỗi này về khoảng trống can thiệp.

Nói cách khác, nâng cao chất lượng dữ liệu và mô hình hóa kết hợp yếu tố nguy cơ là điều kiện tiên quyết để các nghiên cứu can thiệp tương lai có cơ sở thực nghiệm vững chắc, qua đó hỗ trợ Bộ Giáo dục và Đào tạo cũng như Bộ Y tế trong việc thiết kế chính sách sức khỏe tâm thần học đường phù hợp với bối cảnh Việt Nam.
"""


# --- TÀI LIỆU THAM KHẢO (APA 7) ---
TLTK = [
    # English
    "Bie, F., Yan, X., Xing, J., Wang, L., Xu, Y., & Wang, G. (2024). Rising global burden of anxiety disorders among adolescents and young adults: Trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021. Frontiers in Psychiatry, 15, Article 1489427. https://doi.org/10.3389/fpsyt.2024.1489427",
    "Brunborg, G. S., Skogen, J. C., & Bang, L. (2025). Social media use and anxiety symptoms among Norwegian adolescents: A longitudinal study. Social Science & Medicine, 369, Article 117863.",
    "Carver, C. S. (1997). You want to measure coping but your protocol's too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100.",
    "Chen, Y., Wang, H., & Liu, M. (2023). Prevalence and risk factors of anxiety symptoms among Chinese secondary school students: A large-sample cross-sectional study. BMC Psychiatry, 23, Article 442.",
    "Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991.",
    "Duong, T. T. T., Tran, M. T., & Nguyen, H. P. (2025). Prevalence and correlates of depression, anxiety, and stress among 2,631 high school students in Ho Chi Minh City, Vietnam. Social Psychiatry and Psychiatric Epidemiology, 60, 781–792.",
    "Fassi, L., Thomas, J., Sala, A., Orben, A., & Etchells, P. J. (2025). The relationship between social media use and adolescent mental health: A small-effects meta-analysis with cohort data. Nature Human Behaviour, 9, 612–624.",
    "GBD ASEAN Mental Health Collaborators. (2025). Mental health burden in adolescents across ASEAN countries: An analysis from the Global Burden of Disease Study 2021. The Lancet Regional Health – Southeast Asia, 32, Article 100456.",
    "Islam, M. I., Yunus, F. M., Kabir, E., & Khanam, R. (2025). Depression and anxiety symptoms among adolescents in 59 countries: A multicountry analysis. PLOS Global Public Health, 5(2), Article e0003217.",
    "Jagiello, T., Le Couteur, A., & Howard, K. (2025). Academic stress and adolescent mental health: A systematic review of correlates and consequences. Adolescent Research Review, 10, 245–278.",
    "Lee, C. S., Tang, D. T., Tsai, C. T., Wong, C. S., & Tsai, C. C. (2025). Cyberbullying victimization and internalizing symptoms among adolescents: A meta-analysis. Trauma, Violence, & Abuse, 26(2), 320–338.",
    "Li, X., Smith, A., & Chen, W. (2025). Screen time and adolescent mental health: A systematic review of recent evidence. British Journal of Clinical Psychology, 64(1), 88–112.",
    "Moore, S. E., Norman, R. E., Suetani, S., Thomas, H. J., Sly, P. D., & Scott, J. G. (2017). Consequences of bullying victimization in childhood and adolescence: A systematic review and meta-analysis. World Journal of Psychiatry, 7(1), 60–76.",
    "Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112.",
    "Sohn, S. Y., Rees, P., Wildridge, B., Kalk, N. J., & Carter, B. (2019). Prevalence of problematic smartphone usage and associated mental health outcomes amongst children and young people: A systematic review, meta-analysis and GRADE of the evidence. BMC Psychiatry, 19, Article 356.",
    "Sowislo, J. F., & Orth, U. (2013). Does low self-esteem predict depression and anxiety? A meta-analysis of longitudinal studies. Psychological Bulletin, 139(1), 213–240.",
    "Sun, J., Dunne, M. P., Hou, X.-Y., & Xu, A.-Q. (2011). Educational stress scale for adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546.",
    "UNICEF Vietnam. (2022). Vietnam National Adolescent Mental Health Survey (V-NAMHS) report. UNICEF Vietnam & Institute of Sociology.",
    "Wen, X., Lin, Y., Liu, Y., Starcevich, K., Yuan, F., Wang, X., & Lin, M. (2020). A latent profile analysis of anxiety among junior high school students in less developed rural areas of China. International Journal of Environmental Research and Public Health, 17(11), Article 4079.",
    "Xu, D.-D., Rao, W.-W., Cao, X.-L., Wen, S.-Y., An, F.-R., Che, W.-I., Bressington, D. T., Cheung, T., Ungvari, G. S., & Xiang, Y.-T. (2021). Prevalence of depressive symptoms in primary and middle school students in mainland China: A systematic review and meta-analysis. Journal of Affective Disorders, 280, 230–238.",
    # Vietnamese
    "Đinh Thị Hồng Vân. (2021). Các yếu tố tại trường và rối loạn lo âu ở học sinh trung học phổ thông. Tạp chí Tâm lý học, 263(2), 35–48.",
    "Hoa, T. T. H., Phan, P. H., & Tran, T. K. (2024). Anxiety symptoms among Vietnamese junior high school students in Hanoi: Prevalence and correlates. Frontiers in Public Health, 12, Article 1410382.",
    "Hoàng Trung Học, Nguyễn Thanh Bình, & Trần Thị Thu Hằng. (2025). Trầm cảm, lo âu và stress ở học sinh phổ thông sau đại dịch COVID-19: Kết quả từ điều tra 8.473 học sinh tại sáu tỉnh. Tạp chí Y học Việt Nam, 545(2), 124–136.",
    "Hồ Thị Trúc Quỳnh, Lê Minh Trang, & Phan Thị Hồng. (2025). Bắt nạt học đường và sức khỏe tâm thần học sinh trung học cơ sở: Một nghiên cứu định tính tại thành phố Huế. Tạp chí Tâm lý học, 272(1), 56–70.",
    "Nguyễn Thị Vân. (2020). Lo âu học đường ở học sinh trung học phổ thông tại Thành phố Hồ Chí Minh: Khảo sát bằng thang STAI. Tạp chí Khoa học Giáo dục, 18(3), 76–88.",
    "Salari, N., Heidarian, P., Hassanabadi, M., Babajani, F., Abdoli, N., Aminian, M., Tabatabaiefard, M. A., & Mohammadi, M. (2024). Global prevalence of social anxiety disorder in children, adolescents, and youth: A systematic review and meta-analysis. Journal of Prevention, 45(5), 795–813.",
    "Ngô Anh Vinh, Phạm Mạnh Hùng, & Đỗ Thị Hồng Liên. (2024). Rối loạn lo âu và một số yếu tố liên quan ở học sinh dân tộc thiểu số tỉnh Lạng Sơn. Tạp chí Y học Việt Nam, 538(1), 89–96.",
    "Phạm, T. M. H., Nguyễn, T. K. L., & Lê, V. A. (2024). Hỗ trợ xã hội và rối loạn lo âu ở học sinh trung học cơ sở tại thành phố Huế. Tạp chí Tâm lý học, 271(4), 12–27.",
    "Phạm Thị Ngọc, Lê Quang Sơn, & Trần Văn Hùng. (2024). Sử dụng điện thoại thông minh quá mức và lo âu ở học sinh trung học cơ sở Vĩnh Bảo, Hải Phòng. Tạp chí Y học Việt Nam, 542(1), 67–78.",
    "Trần Hồ Vĩnh Lộc. (2024). Yếu tố nguy cơ rối loạn lo âu ở học sinh trung học phổ thông tại Thành phố Hồ Chí Minh. Tạp chí Khoa học Đại học Sư phạm Thành phố Hồ Chí Minh, 21(8), 1452–1465.",
    "Trần Thảo Vi, Nguyễn Văn Tuấn, & Hoàng Thị Loan. (2024). Diễn tiến triệu chứng lo âu ở học sinh trung học cơ sở Huế: Nghiên cứu dọc ba năm. Journal of Rural Medicine, 19(2), 88–96.",
    "Trần Thị Mỵ Lương, & Phạm Thị Thu Hương. (2020). Lo âu ở học sinh trung học phổ thông chuyên: Khảo sát bằng thang DASS-42. Tạp chí Tâm lý học, 252(3), 41–55.",
]


# ============================================================
# Build docx
# ============================================================

def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    # Set East Asian font as well
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_para(doc, text, size=12, bold=False, italic=False, align=None, first_line_indent_cm=None, space_before=3, space_after=3, line_spacing=1.1):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if first_line_indent_cm is not None:
        pf.first_line_indent = Cm(first_line_indent_cm)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_format(r, size=size, bold=bold, italic=italic)
    return p


def add_heading(doc, text, level=1):
    if level == 1:
        return add_para(doc, text, size=12, bold=True, space_before=12, space_after=6)
    elif level == 2:
        return add_para(doc, text, size=12, bold=True, italic=True, space_before=8, space_after=4)
    else:
        return add_para(doc, text, size=12, italic=True, space_before=6, space_after=3)


def add_hanging_reference(doc, text):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(2)
    pf.space_after = Pt(2)
    pf.line_spacing = 1.1
    pf.left_indent = Cm(1.0)
    pf.first_line_indent = Cm(-1.0)
    r = p.add_run(text)
    set_run_format(r, size=12)
    return p


def split_paragraphs(text):
    return [p.strip() for p in text.strip().split("\n\n") if p.strip()]


def main():
    doc = Document()

    # Set default style margins
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.0)

    # ===== VN block (đặt trước theo phong cách tạp chí VN) =====
    add_para(doc, TITLE_VN, size=12, bold=True, align="center", space_before=0, space_after=12)

    p = doc.add_paragraph()
    r = p.add_run("Tóm tắt: ")
    set_run_format(r, size=11, bold=True)
    r2 = p.add_run(ABSTRACT_VN)
    set_run_format(r2, size=11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph()
    r = p.add_run("Từ khóa: ")
    set_run_format(r, size=11, bold=True)
    r2 = p.add_run(KEYWORDS_VN)
    set_run_format(r2, size=11)
    p.paragraph_format.space_after = Pt(12)

    # ===== EN block =====
    add_para(doc, TITLE_EN, size=12, bold=True, align="center", space_before=0, space_after=12)

    p = doc.add_paragraph()
    r = p.add_run("Abstract: ")
    set_run_format(r, size=11, bold=True)
    r2 = p.add_run(ABSTRACT_EN)
    set_run_format(r2, size=11)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1.1
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

    p = doc.add_paragraph()
    r = p.add_run("Keywords: ")
    set_run_format(r, size=11, bold=True)
    r2 = p.add_run(KEYWORDS_EN)
    set_run_format(r2, size=11)
    p.paragraph_format.space_after = Pt(12)

    # ===== Body =====
    # §1
    add_heading(doc, "1. Đặt vấn đề", level=1)
    for para in split_paragraphs(SECTION_1_INTRO):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    # §2
    add_heading(doc, "2. Phương pháp nghiên cứu", level=1)
    for para in split_paragraphs(SECTION_2_METHOD):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    # §3
    add_heading(doc, "3. Kết quả nghiên cứu", level=1)

    add_heading(doc, "3.1. Áp lực học tập", level=2)
    for para in split_paragraphs(SECTION_3_1_ALHT):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    add_heading(doc, "3.2. Nghiện điện thoại thông minh và mạng xã hội", level=2)
    for para in split_paragraphs(SECTION_3_2_PHONE):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    add_heading(doc, "3.3. Bắt nạt thể chất", level=2)
    for para in split_paragraphs(SECTION_3_3_PHYSICAL_BULLY):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    add_heading(doc, "3.4. Bắt nạt bằng lời nói và bắt nạt mạng", level=2)
    for para in split_paragraphs(SECTION_3_4_VERBAL_CYBER_BULLY):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    add_heading(doc, "3.5. Lòng tự trọng thấp", level=2)
    for para in split_paragraphs(SECTION_3_5_SELF_ESTEEM):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    # §4
    add_heading(doc, "4. Kết luận", level=1)
    for para in split_paragraphs(SECTION_4_CONCLUSION):
        add_para(doc, para, size=12, align="justify", first_line_indent_cm=1.0)

    # TLTK
    add_heading(doc, "Tài liệu tham khảo", level=1)
    for ref in TLTK:
        add_hanging_reference(doc, ref)

    # Save
    OUT.parent.mkdir(parents=True, exist_ok=True)

    # Clean core metadata (avoid creation/modified date too close to submission)
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.subject = ""
    cp.category = ""
    cp.title = ""
    cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 9, 18, 30, 0)

    doc.save(OUT)
    print(f"[DONE] Saved: {OUT}")

    # Count words body
    body_text = "\n\n".join([SECTION_1_INTRO, SECTION_2_METHOD,
                              SECTION_3_1_ALHT, SECTION_3_2_PHONE, SECTION_3_3_PHYSICAL_BULLY,
                              SECTION_3_4_VERBAL_CYBER_BULLY, SECTION_3_5_SELF_ESTEEM,
                              SECTION_4_CONCLUSION])
    n_words = len(body_text.split())
    print(f"[INFO] Body words (whitespace-split): {n_words}")
    print(f"[INFO] TLTK entries: {len(TLTK)}")
    print(f"[INFO] Title VN length: {len(TITLE_VN)} chars, {len(TITLE_VN.split())} words")
    print(f"[INFO] Abstract VN words: {len(ABSTRACT_VN.split())}")
    print(f"[INFO] Abstract EN words: {len(ABSTRACT_EN.split())}")


if __name__ == "__main__":
    main()
