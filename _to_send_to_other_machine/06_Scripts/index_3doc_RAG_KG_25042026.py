# -*- coding: utf-8 -*-
"""Index 3 doc dịch + phản biện vào RAG + KG.
Bước:
1. Chunking 3 docx theo section → ~30 chunks
2. Add to RAG rag_db_full (collection lo_au_full)
3. Add nodes + edges vào KG (06_Scripts/kg_data/nodes.json + edges.json)
4. Verification queries
"""
import sys, io, os, json, re
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document

DOC_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien'
KG_DIR  = r'c:\Users\OS\OneDrive\read_books\Lo-au\06_Scripts\kg_data'
RAG_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\rag_dich_phan_bien'
COLLECTION = 'lo_au_dich_phan_bien'  # tách riêng vì rag_db_full bị OneDrive sync conflict

DOCS = [
    {
        'file': 'B5_Brown_Carter_2025_dich_phan_bien_25042026.docx',
        'paper_id': 'QT042_B5',
        'descriptor': 'Brown_Carter_2025_J_Mental_Health_Editorial_UK_school_interventions',
        'authors': 'Brown JSL, Carter B (2025)',
        'journal': 'J Mental Health 34(4):357–361',
        'doi': '10.1080/09638237.2025.2512332',
        'type_paper': 'editorial',
        'country': 'UK',
        'topic': 'school_interventions_review',
    },
    {
        'file': 'PLACES_Brown_2022_dich_phan_bien_25042026.docx',
        'paper_id': 'QT042_PLACES',
        'descriptor': 'Brown_2022_IJERPH_PLACES_Model_help_seeking',
        'authors': 'Brown JSL, Lisk S, Carter B, Stevelink SAM, Van Lieshout R, Michelson D (2022)',
        'journal': 'IJERPH 19(5):2831',
        'doi': '10.3390/ijerph19052831',
        'pmc': 'PMC8909998',
        'type_paper': 'model_paper',
        'country': 'UK_Canada',
        'topic': 'PLACES_help_seeking_model',
    },
    {
        'file': 'BESST_Brown_2024_dich_phan_bien_25042026.docx',
        'paper_id': 'QT042_BESST',
        'descriptor': 'Brown_2024_Lancet_Psychiatry_BESST_cluster_RCT_DISCOVER',
        'authors': 'Brown J, James K, Lisk S, Shearer J, Byford S, Stallard P, Deighton J, '
                   'Saunders D, Yarrum J, Fonagy P, Weaver T, Sclare I, Day C, Evans C, Carter B (2024)',
        'journal': 'Lancet Psychiatry 11(7):504-515',
        'doi': '10.1016/S2215-0366(24)00101-9',
        'pmid': '38759665',
        'isrctn': 'ISRCTN90912799',
        'type_paper': 'cluster_RCT',
        'country': 'UK',
        'topic': 'BESST_DISCOVER_school_CBT_RCT',
        'n_total': 900, 'n_schools': 57, 'n_regions': 4,
        'effect_d_overall': -0.17, 'effect_d_subgroup_elevated': -0.52,
        'icer_per_qaly': 15387,
    },
]

# ============================================================
# BƯỚC 1 — CHUNKING
# ============================================================
print('='*60)
print('BƯỚC 1 — CHUNKING 3 DOC THEO SECTION')
print('='*60)

def docx_to_sections(path):
    """Chia doc thành chunks theo H1/H2 + tóm tắt từng đoạn."""
    d = Document(path)
    chunks = []
    current_h = ''
    current_text = []
    for p in d.paragraphs:
        # detect heading by font size + bold
        is_heading = False
        for r in p.runs:
            if r.bold and r.font.size and r.font.size.pt >= 13:
                is_heading = True
                break
        text = p.text.strip()
        if not text: continue
        if is_heading and (text.startswith('PHẦN') or re.match(r'^\d+\. ', text) or
                           text.startswith('THÔNG TIN') or text.startswith('KẾT THÚC') or
                           text.isupper() or 'Bảng' in text[:10]):
            # flush previous
            if current_text:
                chunks.append((current_h, '\n'.join(current_text)))
            current_h = text
            current_text = [text]
        else:
            current_text.append(text)
    # flush last
    if current_text:
        chunks.append((current_h, '\n'.join(current_text)))
    # also extract tables as separate chunks
    for ti, t in enumerate(d.tables):
        rows_txt = []
        for row in t.rows:
            rows_txt.append(' | '.join(c.text.strip() for c in row.cells))
        chunks.append((f'BẢNG {ti+1}', '\n'.join(rows_txt)))
    # Filter: keep only chunks with substantial content (>200 chars)
    chunks = [(h, t) for h, t in chunks if len(t) > 200]
    return chunks

all_chunks = []
for doc_meta in DOCS:
    path = os.path.join(DOC_DIR, doc_meta['file'])
    chunks = docx_to_sections(path)
    print(f'  {doc_meta["paper_id"]}: {len(chunks)} chunks')
    for hi, (h, t) in enumerate(chunks):
        all_chunks.append({
            'chunk_id': f'{doc_meta["paper_id"]}_chunk_{hi:02d}',
            'paper_id': doc_meta['paper_id'],
            'section': h[:80],
            'text': t[:3000],  # limit chunk size
            'metadata': doc_meta,
        })
print(f'  TỔNG: {len(all_chunks)} chunks')

# ============================================================
# BƯỚC 2 — ADD TO RAG
# ============================================================
print()
print('='*60)
print('BƯỚC 2 — ADD CHUNKS TO RAG (rag_db_full / lo_au_full)')
print('='*60)

os.environ['HF_HOME'] = r'D:\hf_cache'
import chromadb
from sentence_transformers import SentenceTransformer
model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
client = chromadb.PersistentClient(path=RAG_DIR)
try:
    col = client.get_collection(COLLECTION)
    print(f'  Collection {COLLECTION} đã tồn tại với {col.count()} chunks')
except Exception:
    col = client.create_collection(COLLECTION, metadata={'hnsw:space': 'cosine'})
    print(f'  Tạo collection mới {COLLECTION}')
print(f'  Trước: collection có {col.count()} chunks')

# Encode + add
ids = [c['chunk_id'] for c in all_chunks]
docs_text = []
metas = []
for c in all_chunks:
    full_text = (
        f"PAPER {c['paper_id']} — {c['metadata']['descriptor']}\n"
        f"Authors: {c['metadata']['authors']}\n"
        f"Journal: {c['metadata']['journal']}\n"
        f"Type: {c['metadata']['type_paper']} | Country: {c['metadata']['country']}\n"
        f"--- SECTION: {c['section']} ---\n"
        f"{c['text']}"
    )
    docs_text.append(full_text)
    # Chroma metadata only accepts str/int/float/bool
    md = {
        'paper_id': c['paper_id'],
        'section': c['section'],
        'authors': c['metadata']['authors'][:200],
        'journal': c['metadata']['journal'],
        'doi': c['metadata'].get('doi', ''),
        'type_paper': c['metadata']['type_paper'],
        'country': c['metadata']['country'],
        'topic': c['metadata']['topic'],
        'source': 'B5_PLACES_BESST_dich_phan_bien_25042026',
    }
    metas.append(md)

# Encode
print(f'  Encoding {len(ids)} chunks...')
emb = model.encode(docs_text, show_progress_bar=False).tolist()

# Delete existing chunks if rerunning
existing = col.get(ids=ids)
if existing['ids']:
    col.delete(ids=existing['ids'])
    print(f'  Deleted {len(existing["ids"])} existing chunks (rerun)')

col.add(ids=ids, embeddings=emb, documents=docs_text, metadatas=metas)
print(f'  Sau: collection có {col.count()} chunks (added {len(ids)})')

# ============================================================
# BƯỚC 3 — UPDATE KG
# ============================================================
print()
print('='*60)
print('BƯỚC 3 — UPDATE KG (nodes + edges)')
print('='*60)

with open(os.path.join(KG_DIR, 'nodes.json'), 'r', encoding='utf-8') as f:
    nodes = json.load(f)
with open(os.path.join(KG_DIR, 'edges.json'), 'r', encoding='utf-8') as f:
    edges = json.load(f)
print(f'  Trước: {len(nodes)} nodes, {len(edges)} edges')

new_nodes = []
new_edges = []

for doc_meta in DOCS:
    pid = doc_meta['paper_id']
    # Skip if exists
    if any(n.get('id') == pid for n in nodes): continue
    new_nodes.append({
        'id': pid,
        'type': 'Paper',
        'descriptor': doc_meta['descriptor'],
        'file': doc_meta['file'],
        'authors': doc_meta['authors'],
        'journal': doc_meta['journal'],
        'doi': doc_meta.get('doi', ''),
        'paper_type': doc_meta['type_paper'],
        'country': doc_meta['country'],
        'topic': doc_meta['topic'],
    })
    # PMID/PMC nodes
    if 'pmid' in doc_meta:
        new_edges.append({'source': pid, 'target': f'PMID::{doc_meta["pmid"]}', 'rel': 'HAS_PMID'})
    if 'pmc' in doc_meta:
        new_edges.append({'source': pid, 'target': f'PMC::{doc_meta["pmc"]}', 'rel': 'HAS_PMC'})
    if 'isrctn' in doc_meta:
        new_edges.append({'source': pid, 'target': f'TrialReg::{doc_meta["isrctn"]}', 'rel': 'TRIAL_REG'})
    # Country
    new_edges.append({'source': pid, 'target': f'Country::{doc_meta["country"]}', 'rel': 'IN_COUNTRY'})
    # Type
    new_edges.append({'source': pid, 'target': f'StudyType::{doc_meta["type_paper"]}', 'rel': 'HAS_TYPE'})
    # BESST-specific
    if 'n_total' in doc_meta:
        new_edges.append({'source': pid, 'target': f'SampleSize::{doc_meta["n_total"]}', 'rel': 'HAS_N'})
    if 'n_schools' in doc_meta:
        new_edges.append({'source': pid, 'target': f'NumSchools::{doc_meta["n_schools"]}', 'rel': 'HAS_SCHOOLS'})
    if 'effect_d_overall' in doc_meta:
        new_edges.append({'source': pid, 'target': f'EffectSize::Cohen_d::{doc_meta["effect_d_overall"]}', 'rel': 'HAS_EFFECT_OVERALL'})
    if 'effect_d_subgroup_elevated' in doc_meta:
        new_edges.append({'source': pid, 'target': f'EffectSize::Cohen_d::{doc_meta["effect_d_subgroup_elevated"]}', 'rel': 'HAS_EFFECT_SUBGROUP_ELEVATED'})
    if 'icer_per_qaly' in doc_meta:
        new_edges.append({'source': pid, 'target': f'ICER::GBP::{doc_meta["icer_per_qaly"]}', 'rel': 'HAS_ICER'})

# Cross-reference edges (relationships between 3 papers)
new_edges.extend([
    {'source': 'QT042_BESST', 'target': 'QT042_PLACES', 'rel': 'USES_MODEL'},
    {'source': 'QT042_B5', 'target': 'QT042_BESST', 'rel': 'CITES_TRIAL'},
    {'source': 'QT042_B5', 'target': 'QT042_PLACES', 'rel': 'CITES_MODEL'},
    {'source': 'QT042_BESST', 'target': 'Concept::self_referral', 'rel': 'TESTS_CONCEPT'},
    {'source': 'QT042_PLACES', 'target': 'Concept::self_referral', 'rel': 'DEFINES_CONCEPT'},
    {'source': 'QT042_PLACES', 'target': 'Concept::lay_language', 'rel': 'DEFINES_CONCEPT'},
    {'source': 'QT042_PLACES', 'target': 'Concept::publicity', 'rel': 'DEFINES_CONCEPT'},
    {'source': 'QT042_PLACES', 'target': 'Concept::acceptable', 'rel': 'DEFINES_CONCEPT'},
    {'source': 'QT042_PLACES', 'target': 'Concept::convenient', 'rel': 'DEFINES_CONCEPT'},
    {'source': 'QT042_PLACES', 'target': 'Concept::perceived_effectiveness', 'rel': 'DEFINES_CONCEPT'},
    {'source': 'QT042_BESST', 'target': 'Intervention::DISCOVER_CBT_workshop', 'rel': 'TESTS_INTERVENTION'},
    {'source': 'QT042_BESST', 'target': 'Outcome::MFQ_depression', 'rel': 'PRIMARY_OUTCOME'},
    {'source': 'QT042_BESST', 'target': 'Outcome::RCADS_anxiety', 'rel': 'SECONDARY_OUTCOME'},
    {'source': 'QT042_BESST', 'target': 'Outcome::WEMWBS_wellbeing', 'rel': 'SECONDARY_OUTCOME'},
    {'source': 'QT042_BESST', 'target': 'Outcome::SCI_sleep', 'rel': 'SECONDARY_OUTCOME'},
    {'source': 'QT042_BESST', 'target': 'Outcome::CYRM12_resilience', 'rel': 'SECONDARY_OUTCOME'},
    {'source': 'QT042_BESST', 'target': 'Provider::MHST', 'rel': 'DELIVERED_BY'},
    {'source': 'QT042_BESST', 'target': 'AgeGroup::16_18', 'rel': 'TARGET_AGE'},
    # Cross-references with existing corpus
    {'source': 'QT042_BESST', 'target': 'QT08', 'rel': 'CONTRASTS_WITH_LPA_China'},  # Wen 2020
    {'source': 'QT042_BESST', 'target': 'QT29', 'rel': 'CONSISTENT_WITH_NMA_CBT'},   # Li 2025
])

nodes.extend(new_nodes)
edges.extend(new_edges)

with open(os.path.join(KG_DIR, 'nodes.json'), 'w', encoding='utf-8') as f:
    json.dump(nodes, f, ensure_ascii=False, indent=2)
with open(os.path.join(KG_DIR, 'edges.json'), 'w', encoding='utf-8') as f:
    json.dump(edges, f, ensure_ascii=False, indent=2)

print(f'  Sau: {len(nodes)} nodes (+{len(new_nodes)}), {len(edges)} edges (+{len(new_edges)})')

# ============================================================
# BƯỚC 4 — VERIFICATION QUERIES
# ============================================================
print()
print('='*60)
print('BƯỚC 4 — VERIFICATION QUERIES (RAG + KG)')
print('='*60)

print()
print('--- RAG queries ---')
queries = [
    ('BESST trial Cohen d effect size primary outcome', 'QT042_BESST'),
    ('PLACES model 6 thành phần Publicity Lay Acceptable', 'QT042_PLACES'),
    ('BESST self-referral 80% non-consulters MHST', 'QT042_BESST'),
    ('Editorial Brown Carter 2025 universal targeted MYRIAD', 'QT042_B5'),
    ('cost-effectiveness ICER £15387 QALY NICE threshold', 'QT042_BESST'),
    ('Brown 2022 PLACES depression workshop self-confidence 28 120', 'QT042_PLACES'),
]
for q, expected_id in queries:
    emb_q = model.encode([q]).tolist()
    r = col.query(query_embeddings=emb_q, n_results=3, include=['metadatas','distances','documents'])
    print(f'\\n  Q: \"{q[:55]}\"')
    print(f'  Expected: {expected_id}')
    for i, (m, dist) in enumerate(zip(r['metadatas'][0], r['distances'][0])):
        match = '✓' if m['paper_id'] == expected_id else ' '
        print(f'    {match} #{i+1} [{dist:.3f}] {m["paper_id"]} | {m["section"][:60]}')

print()
print('--- KG queries ---')
# Find all edges from QT042_BESST
besst_edges = [e for e in edges if e['source'] == 'QT042_BESST']
places_edges = [e for e in edges if e['source'] == 'QT042_PLACES']
b5_edges = [e for e in edges if e['source'] == 'QT042_B5']
print(f'  QT042_BESST có {len(besst_edges)} outgoing edges')
print(f'  QT042_PLACES có {len(places_edges)} outgoing edges')
print(f'  QT042_B5 có {len(b5_edges)} outgoing edges')
print('  Sample BESST edges:')
for e in besst_edges[:8]:
    print(f'    {e["source"]} --[{e["rel"]}]--> {e["target"]}')

# Find concepts shared between BESST and PLACES
besst_targets = {e['target'] for e in besst_edges}
places_targets = {e['target'] for e in places_edges}
shared = besst_targets & places_targets
print(f'\\n  Concepts/targets shared BESST <-> PLACES: {sorted(shared)}')

print()
print('='*60)
print('HOÀN TẤT INDEX + VERIFICATION')
print('='*60)
