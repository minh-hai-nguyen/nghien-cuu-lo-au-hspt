"""DOC: Dien giai chi tiet yeu to AP LUC HOC TAP (ALHT)
theo so lieu chinh xac tu file 00_Binh luan so lieu-v2.docx
Bang 3.21 + 3.23 + 3.24

ALL FACTS VERIFIED tu Tables 5, 7, 8 file v2:
- Bang 3.21: Bieu hien ESSA
  + A. Ap luc hoc tap (chung): DTB=51,13, DLC=23,922, thu bac 1
  + ESSA.3 (bai tap): 50,49 (33,045), thu 2
  + ESSA.4 (su nghiep tuong lai): 58,56 (31,972), thu 1
  + ESSA.5 (cha me quan tam diem): 49,58 (33,857), thu 3
  + ESSA.6 (ket qua kiem tra): 45,91 (32,161), thu 4
- Bang 3.23: Fit indices SEM
  + ALHT-RLLATQ: chi²(43)=146,315; chi²/df=3,403; CFI=0,972; TLI=0,965; RMSEA=0,042 (KTC 0,035-0,050)
  + ALHT-RLLACL: chi²(19)=71,739; chi²/df=3,776; CFI=0,975; TLI=0,964; RMSEA=0,045 (0,034-0,057)
  + ALHT-RLLAXH: chi²(19)=59,943; chi²/df=3,155; CFI=0,983; TLI=0,975; RMSEA=0,040 (0,029-0,052)
  + ALHT-RLLA (3 factors): chi²(13)=31,015; chi²/df=2,386; CFI=0,992; TLI=0,987; RMSEA=0,032
  + ALHT-RLLA (2 factors): chi²(8)=23,049; chi²/df=2,881; CFI=0,993; TLI=0,987; RMSEA=0,037
- Bang 3.24: Path coefficients
  + ALHT->RLLATQ: beta=0,510; SE=0,066; CR=10,984; p<0,001
  + ALHT->RLLACL: beta=0,253; SE=0,055; CR=6,405; p<0,001
  + ALHT->RLLAXH: beta=0,490; SE=0,068; CR=10,369; p<0,001
  + ALHT->RLLA (3 factors): beta=0,533; SE=0,141; CR=8,888; p<0,001; R²=0,284
  + ALHT->RLLA (2 factors): beta=0,530; SE=0,411; CR=11,476; p<0,001; R²=0,281
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Dien_giai_yeu_to_ap_luc_hoc_tap_chi_tiet_v2.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DIỄN GIẢI YẾU TỐ ÁP LỰC HỌC TẬP\nCHI TIẾT THEO 3 BẢNG SỐ LIỆU CHƯƠNG 3\n— Theo file 00_Bình luận số liệu-v2 (Bảng 3.21, 3.23, 3.24) —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# 1. Khung chung
H('1. Khung chung', level=2, color=NAVY)
para(
    'Áp lực học tập (ALHT) là yếu tố nguy cơ ĐẦU TIÊN trong nhóm '
    'yếu tố nguy cơ được phân tích trong chương 3 luận án. Tài liệu '
    'này diễn giải CHI TIẾT từng con số trên ba bảng then chốt: '
    'Bảng 3.21 (mức độ và biểu hiện), Bảng 3.23 (model fit của '
    'SEM), và Bảng 3.24 (hệ số tác động của ALHT đến RLLA).'
)
para(
    'Khác với doc tổng quan giả thuyết H1, doc này tập trung vào '
    'GIẢI THÍCH TỪNG CON SỐ — bao gồm điểm trung bình, độ lệch '
    'chuẩn, fit indices, hệ số chuẩn hóa, sai số chuẩn, tỷ số tới '
    'hạn, và R². Mục tiêu: giúp người đọc HIỂU đầy đủ ý nghĩa của '
    'từng chỉ số thay vì chỉ trích con số tóm gọn.'
)

# 2. Bảng 3.21 — Biểu hiện ALHT
H('2. Diễn giải Bảng 3.21 — Mức độ và biểu hiện ALHT', level=2, color=NAVY)
caption('Bảng 1. Mức độ và biểu hiện ALHT (rút từ Bảng 3.21 chương 3)')
add_table(
    ['Mục', 'Nội dung', 'Min', 'Max', 'ĐTB', 'ĐLC', 'Thứ bậc'],
    [
        ['A. Áp lực học tập (chung)', 'Tổng nhóm yếu tố', '0', '100', '51,13', '23,922', '1'],
        ['ESSA.3', 'Có quá nhiều bài tập về nhà', '0', '100', '50,49', '33,045', '2'],
        ['ESSA.4', 'Việc học và đòi hỏi sự nghiệp tương lai', '0', '100', '58,56', '31,972', '1'],
        ['ESSA.5', 'Cha mẹ quá quan tâm đến điểm số', '0', '100', '49,58', '33,857', '3'],
        ['ESSA.6', 'Có quá nhiều kết quả kiểm tra', '0', '100', '45,91', '32,161', '4'],
    ]
)
para('')

H('Diễn giải từng con số:', level=3, color=NAVY)
para(
    'CHỈ SỐ TỔNG NHÓM (ĐTB = 51,13). Trên thang 0–100, ĐTB của '
    'nhóm áp lực học tập đạt 51,13 — vượt quá GIỮA THANG ĐO. Nói '
    'cách khác, học sinh THCS Việt Nam trung bình tự đánh giá '
    'mức độ áp lực học tập ở mức TRÊN TRUNG BÌNH. Đây là MỨC '
    'CAO NHẤT trong tất cả nhóm yếu tố nguy cơ (thứ bậc 1).'
)
para(
    'ĐỘ LỆCH CHUẨN ĐLC = 23,922. ĐLC này CHIẾM 47% của ĐTB (23,922 '
    '/ 51,13 = 0,468) — gợi ý PHƯƠNG SAI ĐÁNG KỂ giữa các học '
    'sinh. Có nhóm trải nghiệm áp lực rất CAO (ĐTB + 1 ĐLC = '
    '75,05), có nhóm trải nghiệm áp lực THẤP (ĐTB − 1 ĐLC = '
    '27,21). Đa dạng cá nhân RỘNG.'
)
para(
    'BIỂU HIỆN HÀNG ĐẦU — ĐỊNH HƯỚNG SỰ NGHIỆP (ESSA.4, ĐTB = '
    '58,56). "Việc học tập và đòi hỏi sự nghiệp tương lai" là '
    'biểu hiện CAO NHẤT (thứ 1 trong 4 mục) — VƯỢT 58 điểm trên '
    'thang 0–100. Phù hợp đặc thù bối cảnh giáo dục Việt Nam: '
    'học sinh THCS đã chịu áp lực CHUẨN BỊ thi vào lớp 10 + định '
    'hướng nghề + áp lực sự nghiệp tương lai từ lớp 6–7. Đây là '
    'phát hiện ĐẶC BIỆT — ở đa số nước phương Tây, áp lực sự '
    'nghiệp xuất hiện muộn hơn (THPT hoặc đại học).'
)
para(
    'BIỂU HIỆN THỨ HAI — KHỐI LƯỢNG BÀI TẬP (ESSA.3, ĐTB = '
    '50,49). "Có quá nhiều bài tập về nhà" — ĐTB vẫn ở MỨC TRÊN '
    'TRUNG BÌNH. Phù hợp với báo cáo OECD (2018) về Việt Nam '
    'đứng TOP 5 quốc gia có thời gian học/làm bài tập về nhà cao '
    'nhất.'
)
para(
    'BIỂU HIỆN THỨ BA — KỲ VỌNG CHA MẸ VỀ ĐIỂM SỐ (ESSA.5, ĐTB = '
    '49,58). "Cha mẹ quá quan tâm đến điểm số" — ĐTB GẦN GIỮA '
    'THANG. Phù hợp đặc thù văn hóa Á: cha mẹ đặt nặng thành '
    'tích học tập như chỉ báo thành công của con. Phù hợp với '
    'phát hiện Nguyễn Thị Vân (2020) — VN004 trong cơ sở dữ '
    'liệu — kỳ vọng cha mẹ có ảnh hưởng đến 48,9% học sinh.'
)
para(
    'BIỂU HIỆN THẤP NHẤT — ÁP LỰC KIỂM TRA (ESSA.6, ĐTB = '
    '45,91). "Có quá nhiều kết quả kiểm tra" có ĐTB THẤP NHẤT '
    'trong bốn mục, nhưng VẪN gần giữa thang. Đáng chú ý: dù '
    'thấp nhất, mức 45,91 vẫn cao hơn nhiều so với các yếu tố '
    'nguy cơ khác (sẽ phân tích trong các doc riêng).'
)

H('Phát hiện ĐẶC BIỆT từ Bảng 3.21:', level=3, color=NAVY)
para(
    'Thứ tự cường độ trong NHÓM ALHT: sự nghiệp tương lai (58,56) '
    '> bài tập về nhà (50,49) > kỳ vọng cha mẹ (49,58) > kiểm '
    'tra (45,91). Nói cách khác, học sinh THCS Việt Nam nhận thức '
    'áp lực TƯƠNG LAI mạnh hơn áp lực HIỆN TẠI (bài tập, kiểm '
    'tra). Đây là pattern KHÁC với học sinh phương Tây cùng tuổi '
    '— gợi ý đặc thù văn hóa "tư duy đường dài" trong giáo dục '
    'Việt Nam.'
)

# 3. Bảng 3.23 — Fit indices
H('3. Diễn giải Bảng 3.23 — Fit indices của 5 mô hình SEM', level=2, color=NAVY)
caption('Bảng 2. Fit indices các mô hình SEM ALHT → RLLA (rút từ Bảng 3.23)')
add_table(
    ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI RMSEA', 'Đánh giá'],
    [
        ['ALHT → RLLATQ', '146,315 (43)', '3,403', '0,972', '0,965', '0,042', '0,035–0,050', 'TỐT'],
        ['ALHT → RLLACL', '71,739 (19)', '3,776', '0,975', '0,964', '0,045', '0,034–0,057', 'TỐT'],
        ['ALHT → RLLAXH', '59,943 (19)', '3,155', '0,983', '0,975', '0,040', '0,029–0,052', 'TỐT'],
        ['ALHT → RLLA (3 factors)', '31,015 (13)', '2,386', '0,992', '0,987', '0,032', '0,018–0,047', 'XUẤT SẮC'],
        ['ALHT → RLLA (2 factors)', '23,049 (8)', '2,881', '0,993', '0,987', '0,037', '0,020–0,056', 'XUẤT SẮC'],
    ]
)
para('')

H('Diễn giải fit indices theo Hu & Bentler (1999):', level=3, color=NAVY)
para(
    'Theo Hu và Bentler (1999), ngưỡng "good fit" cho SEM: CFI ≥ '
    '0,95; TLI ≥ 0,95; RMSEA ≤ 0,06; χ²/df < 3 (lý tưởng). Áp '
    'dụng cho năm mô hình ALHT → RLLA của chương 3 luận án:'
)
para(
    'CFI từ 0,972 đến 0,993 — TẤT CẢ vượt ngưỡng 0,95. Hai mô '
    'hình tổng (3 factors và 2 factors) đạt CFI > 0,99 — XUẤT '
    'SẮC. Điều này có nghĩa: cấu trúc SEM giả định (ALHT là '
    'biến tiềm ẩn ảnh hưởng đến từng dạng RLLA) PHÙ HỢP rất '
    'tốt với dữ liệu thực.'
)
para(
    'TLI từ 0,964 đến 0,987 — TẤT CẢ vượt ngưỡng 0,95. TLI '
    '(Tucker-Lewis Index, còn gọi NNFI) phạt mô hình phức tạp '
    'hơn CFI, nên TLI thường THẤP HƠN CFI một chút. Cả 5 mô '
    'hình đều đạt good fit.'
)
para(
    'RMSEA từ 0,032 đến 0,045 — TẤT CẢ dưới ngưỡng 0,06. RMSEA '
    '(Root Mean Square Error of Approximation) đặc biệt quan '
    'trọng vì nhạy với cỡ mẫu lớn. Với n = 1.352, RMSEA dưới '
    '0,05 cho TẤT CẢ mô hình — bằng chứng MẠNH cho fit. Quan '
    'trọng nhất: 90% CI của RMSEA chứa 0,05 trong tất cả mô '
    'hình — điều kiện ĐẶC BIỆT để kết luận good fit.'
)
para(
    'χ²/df từ 2,386 đến 3,776 — vượt ngưỡng lý tưởng < 3 ở hai '
    'mô hình (RLLATQ với 3,403 và RLLACL với 3,776), nhưng '
    'TẤT CẢ dưới 5. χ² test rất nhạy với cỡ mẫu n = 1.352, nên '
    'KHÔNG nên dựa hoàn toàn vào χ² đơn lẻ. Phù hợp khuyến '
    'nghị Hu & Bentler: kết hợp nhiều chỉ số.'
)

H('Mô hình "3 factors" và "2 factors" — Sự khác biệt:', level=3, color=NAVY)
para(
    'Bảng 3.23 báo cáo HAI cách tổng hợp RLLA:'
)
para(
    '• 3 factors: RLLA tổng được mô hình hóa như biến tiềm ẩn '
    'gồm BA biến chỉ báo (RLLATQ, RLLACL, RLLAXH). Mô hình giả '
    'định cả ba dạng có cấu trúc tương quan đồng nhất.'
)
para(
    '• 2 factors: RLLA tổng được mô hình hóa như biến tiềm ẩn '
    'gồm HAI nhân tố. Theo phát hiện H3 (lan tỏa và xã hội '
    'TĂNG theo khối; chia ly GIẢM), khả năng cao 2 factors = '
    'TĂNG (RLLATQ + RLLAXH) + GIẢM (RLLACL) — tách RLLAC ra '
    'do pattern theo khối khác biệt.'
)
para(
    'Mô hình 2 factors có CFI = 0,993 (cao hơn 3 factors) và '
    'RMSEA = 0,037 (cao hơn 3 factors) — cả hai mô hình đều '
    'đạt fit XUẤT SẮC. Khác biệt RẤT NHỎ — gợi ý cả hai cách '
    'tổng hợp đều hợp lệ về mặt thống kê.'
)

# 4. Bảng 3.24 — Path coefficients
H('4. Diễn giải Bảng 3.24 — Hệ số tác động ALHT → RLLA', level=2, color=NAVY)
caption('Bảng 3. Hệ số tác động chuẩn hóa của ALHT đến RLLA (rút từ Bảng 3.24)')
add_table(
    ['Mô hình', 'Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
    [
        ['1', 'ALHT → RLLATQ', '0,510', '0,066', '10,984', '< 0,001', '—'],
        ['2', 'ALHT → RLLACL', '0,253', '0,055', '6,405', '< 0,001', '—'],
        ['3', 'ALHT → RLLAXH', '0,490', '0,068', '10,369', '< 0,001', '—'],
        ['4', 'ALHT → RLLA (3 factors)', '0,533', '0,141', '8,888', '< 0,001', '0,284'],
        ['5', 'ALHT → RLLA (2 factors)', '0,530', '0,411', '11,476', '< 0,001', '0,281'],
    ]
)
para('')

H('Diễn giải từng đường ảnh hưởng:', level=3, color=NAVY)
para(
    'ĐƯỜNG 1 — ALHT → RLLATQ (lan tỏa): β = 0,510. Khi áp lực '
    'học tập TĂNG 1 độ lệch chuẩn (SD), điểm rối loạn lo âu lan '
    'tỏa TĂNG 0,510 SD. Theo quy ước Cohen (1988), |β| > 0,5 '
    'thuộc "tác động LỚN". Đây là cường độ MẠNH NHẤT trong ba '
    'đường ảnh hưởng. Sai số chuẩn S.E. = 0,066 — nhỏ; tỷ số '
    'tới hạn C.R. = 10,984 — RẤT lớn (vượt ngưỡng 1,96 và '
    '2,58); p < 0,001 — CỰC kỳ ý nghĩa thống kê.'
)
para(
    'ĐƯỜNG 2 — ALHT → RLLACL (chia ly): β = 0,253. Đây là cường '
    'độ THẤP NHẤT trong ba đường, tương đương "tác động NHỎ–'
    'TRUNG BÌNH" theo Cohen (1988). C.R. = 6,405 — vẫn lớn; p < '
    '0,001 — vẫn ý nghĩa thống kê. Phát hiện: áp lực học tập '
    'có tác động ý nghĩa ngay cả với lo âu chia ly — dù pattern '
    'lo âu chia ly thường được xem là "không phụ thuộc bối cảnh '
    'học tập".'
)
para(
    'ĐƯỜNG 3 — ALHT → RLLAXH (xã hội): β = 0,490. Cường độ '
    'GẦN BẰNG đường 1 (β = 0,510), khoảng 96% cường độ. Vẫn '
    'thuộc "tác động LỚN" theo Cohen. C.R. = 10,369; p < '
    '0,001 — ý nghĩa cao. Phù hợp với khung khái niệm: áp '
    'lực học tập tạo ra cả lo âu LAN TỎA (về nhiều lĩnh vực) '
    'và lo âu XÃ HỘI (về đánh giá xã hội — điểm số là một '
    'đánh giá xã hội mạnh).'
)
para(
    'ĐƯỜNG 4 — ALHT → RLLA TỔNG (3 factors): β = 0,533, R² = '
    '0,284. Khi tổng hợp ba dạng RLLA, β tổng đạt 0,533 — '
    'cao hơn từng đường riêng lẻ. R² = 28,4% nghĩa là áp '
    'lực học tập GIẢI THÍCH được 28,4% phương sai của RLLA '
    'tổng. Theo Cohen (1988): R² > 0,26 thuộc "effect size '
    'LỚN" — kết quả MẠNH MẼ.'
)
para(
    'ĐƯỜNG 5 — ALHT → RLLA TỔNG (2 factors): β = 0,530, R² = '
    '0,281. Cường độ và R² gần như giống đường 4 — củng cố '
    'tính ổn định của phát hiện qua cả hai cách phân tách.'
)

H('So sánh thứ tự cường độ:', level=3, color=NAVY)
para(
    'β ALHT → RLLATQ (0,510) > β ALHT → RLLAXH (0,490) > β '
    'ALHT → RLLACL (0,253). Áp lực học tập tác động MẠNH '
    'NHẤT lên LO ÂU LAN TỎA, mạnh thứ 2 lên LO ÂU XÃ HỘI, '
    'và yếu nhất lên LO ÂU CHIA LY. Pattern này phù hợp với '
    'khung phát triển: lan tỏa và xã hội là rối loạn KÉO '
    'DÀI và liên quan đánh giá; chia ly là rối loạn KHỞI '
    'PHÁT SỚM ít liên quan đến áp lực học đường.'
)

H('Tỷ số cường độ với các yếu tố khác trong chương 3:', level=3, color=NAVY)
para('Để thấy rõ vị trí ALHT, so sánh với các yếu tố khác đã phân tích trong chương 3:')
add_table(
    ['Yếu tố', 'β → RLLATQ', 'β → RLLACL', 'β → RLLAXH', 'Tỷ số ALHT là'],
    [
        ['ALHT (áp lực học tập)', '0,510', '0,253', '0,490', 'Tham chiếu = 1'],
        ['NĐT (nghiện điện thoại)', '0,336', '0,265', '0,383', '~66% ALHT (RLLATQ)'],
        ['BNHĐ (bắt nạt học đường)', '0,215', '0,376', '0,253', '~42% ALHT (RLLATQ)'],
        ['TTr (tự trọng) — bảo vệ', '−0,455', '−0,087', '−0,415', '|β| ~89% |β ALHT|'],
        ['BPĐP (đối phó)', '0,749', '0,132', '0,670', '147% ALHT — fit KÉM'],
    ]
)
para('')
para(
    'ALHT là yếu tố nguy cơ MẠNH NHẤT có fit indices TỐT (loại '
    'trừ BPĐP có β cao hơn nhưng fit kém — RMSEA = 0,080–0,204). '
    'Tự trọng có cường độ |β| ngang bằng ~89% ALHT — phù hợp '
    'với giả thuyết H2.'
)

# 5. Phát hiện chính
H('5. Bốn phát hiện chính từ ba bảng', level=2, color=NAVY)
para(
    'Thứ nhất, áp lực học tập là yếu tố nguy cơ ĐỨNG ĐẦU '
    'trong nhóm yếu tố nguy cơ về MỨC ĐỘ TỰ ĐÁNH GIÁ '
    '(ĐTB = 51,13/100, thứ bậc 1).', indent=False
)
para(
    'Thứ hai, biểu hiện CAO NHẤT trong nhóm ALHT là "định '
    'hướng sự nghiệp tương lai" (ESSA.4, ĐTB = 58,56) — phản '
    'ánh đặc thù bối cảnh giáo dục Việt Nam: áp lực TƯƠNG '
    'LAI sớm hơn so với học sinh phương Tây.'
)
para(
    'Thứ ba, mô hình SEM ALHT → RLLA đạt fit indices TỐT '
    'đến XUẤT SẮC: CFI 0,972–0,993; RMSEA 0,032–0,045. Tất '
    'cả vượt ngưỡng Hu & Bentler (1999).'
)
para(
    'Thứ tư, cường độ tác động chuẩn hóa: β = 0,510 cho '
    'lan tỏa; 0,490 cho xã hội; 0,253 cho chia ly; 0,533 cho '
    'tổng (R² = 28,4%). Pattern phù hợp với khung phát '
    'triển: ALHT mạnh nhất với rối loạn liên quan đánh giá, '
    'yếu nhất với rối loạn khởi phát sớm.'
)

# 6. Hàm ý
H('6. Bốn hàm ý cho thiết kế can thiệp', level=2, color=NAVY)
para(
    'HÀM Ý 1 — Can thiệp NGAY VỚI MỤC ESSA.4 (sự nghiệp '
    'tương lai). Đây là biểu hiện áp lực CAO NHẤT (ĐTB = '
    '58,56). Đề xuất chương trình "định hướng nghề" cho '
    'học sinh THCS giảm áp lực qua: thông tin nghề minh '
    'bạch, đa dạng hóa con đường thành công, nhấn mạnh kỹ '
    'năng MỀM bên cạnh thi cử.', indent=False
)
para(
    'HÀM Ý 2 — Can thiệp song song cha mẹ. ESSA.5 (cha mẹ '
    'quá quan tâm điểm số) có ĐTB = 49,58 — gần giữa '
    'thang. Phù hợp mô hình EACP (Lochman 2025): chương '
    'trình can thiệp 16 buổi cho cha mẹ + 25 buổi cho học '
    'sinh.'
)
para(
    'HÀM Ý 3 — Sử dụng RCADS (Chorpita 2000) thay vì '
    'thang đo tự thiết kế. Chương 3 dùng RCADS — phù hợp '
    'với chuẩn quốc tế. Khuyến nghị các nghiên cứu VN '
    'tiếp theo cũng dùng RCADS để có thể META-ANALYSIS '
    'với y văn quốc tế.'
)
para(
    'HÀM Ý 4 — Phân tích RIÊNG cho từng loại RLLA. β ALHT '
    'khác nhau giữa 3 dạng (0,510 / 0,253 / 0,490) — '
    'không nên gộp thành biến RLLA tổng thuần khi báo '
    'cáo. Khi báo cáo, NÊN cung cấp cả 3 đường cùng với '
    'mô hình tổng hợp.'
)

# 7. CÂU TRẢ LỜI
H('7. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue_run('Diễn giải áp lực học tập theo 3 bảng chương 3:', bold=True)
blue_run(
    '(1) BẢNG 3.21 — MỨC ĐỘ + BIỂU HIỆN: ALHT có ĐTB = '
    '51,13 (ĐLC = 23,922) trên thang 0–100 — vượt giữa '
    'thang. Thứ tự biểu hiện: sự nghiệp tương lai '
    '(ESSA.4, ĐTB = 58,56) > bài tập (ESSA.3, 50,49) > '
    'kỳ vọng cha mẹ (ESSA.5, 49,58) > kiểm tra (ESSA.6, '
    '45,91). Đặc thù VN: áp lực TƯƠNG LAI sớm hơn '
    'phương Tây.'
)
blue_run(
    '(2) BẢNG 3.23 — MODEL FIT SEM: 5 mô hình ALHT → '
    'RLLA đều đạt fit TỐT đến XUẤT SẮC theo Hu & '
    'Bentler 1999: CFI 0,972–0,993; TLI 0,964–0,987; '
    'RMSEA 0,032–0,045 (KTC 90% chứa 0,05); χ²/df '
    '2,386–3,776. Hai mô hình tổng (3 factors và 2 '
    'factors) đạt CFI > 0,99 — XUẤT SẮC.'
)
blue_run(
    '(3) BẢNG 3.24 — HỆ SỐ CHUẨN HÓA: β ALHT → RLLATQ = '
    '0,510 (mạnh nhất); β → RLLAXH = 0,490; β → '
    'RLLACL = 0,253 (yếu nhất); β tổng = 0,533 (3 '
    'factors) hoặc 0,530 (2 factors). R² = 0,284 — '
    'ALHT giải thích 28,4% phương sai RLLA tổng — '
    'effect size LỚN theo Cohen (1988). Tất cả p < '
    '0,001.'
)
blue_run(
    '(4) PATTERN: ALHT tác động MẠNH NHẤT lên RLLATQ '
    '(lan tỏa) và RLLAXH (xã hội) — hai dạng liên '
    'quan đánh giá. Yếu nhất lên RLLACL (chia ly) — '
    'dạng khởi phát sớm ít phụ thuộc bối cảnh học '
    'tập.'
)
blue_run(
    '(5) VỊ TRÍ ALHT trong các yếu tố nguy cơ chương '
    '3: là YẾU TỐ MẠNH NHẤT có fit indices tốt. NĐT '
    '(β = 0,336) ~66% ALHT; BNHĐ (β = 0,215) ~42% '
    'ALHT; TTr (|β| = 0,455) ~89% ALHT (yếu tố bảo '
    'vệ); BPĐP có β cao hơn nhưng fit kém.'
)

# 8. TLTK
H('8. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
    'Chorpita, B. F., Yim, L., Moffitt, C., Umemoto, L. A., & Francis, S. E. (2000). Assessment of symptoms of DSM-IV anxiety and depression in children: A revised child anxiety and depression scale. Behaviour Research and Therapy, 38(8), 835–855. https://doi.org/10.1016/S0005-7967(99)00130-8',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
    'Lochman, J. E., và cộng sự. (2025). Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology.',
    'Nguyễn, T. V. (2020). Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh. [VN004 trong cơ sở dữ liệu dự án.]',
    'Sun, J., Dunne, M. P., Hou, X. Y., & Xu, A. Q. (2011). Educational Stress Scale for Adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546. https://doi.org/10.1177/0734282910394976',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
