# -*- coding: utf-8 -*-
"""
(1) Sửa mục Happy House trong Bài 2 v6: thay tên viết tắt người Việt
    bằng TÊN ĐẦY ĐỦ đúng như tạp chí Cambridge Prisms công bố (không bịa).
(2) Sửa Bài 03 (Công Thị Hằng) -> v4: 3 lỗi thư mục + tên tạp chí JAACAP
    + khử chuỗi DOI bị lặp.
File gốc giữ nguyên để đối chứng.
"""
import os
from docx import Document
from datetime import datetime

ROOT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn"
B2_V6 = os.path.join(ROOT, "Bai2_CanThiep_HSTHCS_v6_16052026.docx")
B03_IN = os.path.join(ROOT, "bài-03", "Công Thị Hằng_v3_đã sửa.docx")
B03_OUT = os.path.join(ROOT, "bài-03", "Công Thị Hằng_v4_đã sửa.docx")


def set_para(p, text):
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


# ---- (1) Happy House: tên đầy đủ đúng như Cambridge Prisms in ----
HH_OLD = "Tran, T. D., Nguyen, H., Shochet, I."
HH_NEW = ("Thach Duc Tran, Huong Nguyen, Ian Shochet, Nga Nguyen, Nga La, "
          "Astrid Wurfl, Jayne Orr, Hau Nguyen, Ruby Stocker, & Jane Fisher. "
          "(2023). School-based universal mental health promotion intervention "
          "for adolescents in Vietnam: Two-arm, parallel, controlled trial. "
          "Cambridge Prisms: Global Mental Health, 10, Article e69. "
          "https://doi.org/10.1017/gmh.2023.66")

doc = Document(B2_V6)
done = False
for p in doc.paragraphs:
    if HH_OLD in p.text:
        set_para(p, HH_NEW)
        done = True
        break
print(f"(1) Happy House full names: {'OK' if done else 'NOT FOUND'}")
if done:
    doc.save(B2_V6)
    print(f"    Saved: {B2_V6}")


# ---- (2) Bài 03 ----
doc3 = Document(B03_IN)

# 2a. Sửa cụm con
SUBS = [
    ("children and adolescents. Psychological Bulletin, 140(3)",
     "children and adolescents: 20 years after. Psychological Bulletin, 140(3)"),
    ("Graham, Joseph W. (2009)", "Graham, J. W. (2009)"),
    ("Enders, Craig K. (2010)", "Enders, C. K. (2010)"),
    ("Shek, Daniel T. L. (1997)", "Shek, D. T. L. (1997)"),
    ("Parenting factors associated with depression and anxiety in children: "
     "A systematic review and meta-analysis",
     "Parental factors associated with depression and anxiety in young people: "
     "A systematic review and meta-analysis"),
]
nsub = 0
for sub, new in SUBS:
    for p in doc3.paragraphs:
        if sub in p.text:
            set_para(p, p.text.replace(sub, new))
            nsub += 1
            break
    else:
        print(f"  !! Bài03 SUB NOT FOUND: {sub[:50]}")
print(f"(2a) Bài 03 sửa cụm con: {nsub}/{len(SUBS)}")

# 2b. Sửa tên tạp chí JAACAP ("Child và Adolescent" -> "Child & Adolescent")
njr = 0
for p in doc3.paragraphs:
    if "Child và Adolescent Psychiatry" in p.text:
        set_para(p, p.text.replace("Child và Adolescent Psychiatry",
                                   "Child & Adolescent Psychiatry"))
        njr += 1
print(f"(2b) Bài 03 sửa tên tạp chí JAACAP: {njr} mục")

# 2c. Khử chuỗi DOI bị lặp 2 lần liền nhau
ndoi = 0
for p in doc3.paragraphs:
    t = p.text
    if t.count("https://doi.org/") == 2:
        parts = t.split("https://doi.org/")
        if len(parts) == 3 and parts[1].strip() == parts[2].strip():
            set_para(p, parts[0] + "https://doi.org/" + parts[1])
            ndoi += 1
print(f"(2c) Bài 03 khử DOI lặp: {ndoi} mục")

cp = doc3.core_properties
cp.author = ""
cp.last_modified_by = ""
cp.comments = ""
cp.modified = datetime(2026, 5, 16, 11, 0, 0)
doc3.save(B03_OUT)
print(f"    Saved: {B03_OUT}")
print("[DONE]")
