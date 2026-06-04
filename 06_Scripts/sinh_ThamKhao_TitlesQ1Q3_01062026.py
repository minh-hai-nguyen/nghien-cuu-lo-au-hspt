# -*- coding: utf-8 -*-
"""Sinh doc tham khao titles Q1+Q3 cua tac gia Chau A + Chau Phi
ve roi loan lo au adolescent. De nhom ban chon ten bai.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'ThamKhao_Titles_Q1Q3_AsiaChauPhi_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.25


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=False, size=11):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('▸ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def make_title_table(rows, col_widths_cm):
    """Headers: STT | Title | Authors+Origin | Journal+Q | Notes"""
    t = d.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'; t.autofit = False
    headers = ['#', 'Tiêu đề bài', 'Tác giả + Quốc tịch',
               'Tạp chí + Năm', 'Ghi chú style']
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
H1('THAM KHẢO TIÊU ĐỀ BÀI BÁO Q1 + Q3 VỀ RỐI LOẠN LO ÂU')
P('Bộ sưu tập tiêu đề các bài báo gần đây (2023-2025) của tác giả châu Á '
  'và châu Phi về rối loạn lo âu ở thanh thiếu niên, phục vụ thảo luận '
  'chọn tên cho 2 bài Q1 + Q3 của nhóm.', italic=True)
P('Ngày soạn: 01/06/2026', italic=True)

P('')
P('Nguồn search: PubMed/PMC, BMC Psychiatry, PLOS ONE, PLOS Mental Health, '
  'Journal of Affective Disorders, Frontiers in Psychiatry, Nature Scientific '
  'Reports, General Psychiatry. Tổng cộng ~30+ titles được tổng hợp.',
  italic=True)


# ============================================================
H1('PHẦN 1 — Q1 TIER: SEM / Mô hình tích hợp / Mixed-methods')
P('Các bài Q1 đặc trưng có METHODOLOGICAL CONTRIBUTION (SEM, mediation, '
  'longitudinal, mixed-methods) và đa số là từ 2024-2025.', italic=True)

H2('1.1 Châu Á — Trung Quốc, Việt Nam, Hàn Quốc, Nhật Bản')

q1_asia = [
    ('1', 'Thirty-year trends of anxiety disorders among adolescents based '
     'on the 2019 Global Burden of Disease Study',
     'Liu Y et al. (Trung Quốc)',
     'General Psychiatry (Q1, IF 24)\n2024',
     'GBD-based trend; Title rất dài'),
    ('2', 'Risk factors of depressive and anxiety symptoms in Chinese '
     'adolescents [...]',
     'Wang et al. (Trung Quốc)',
     'Scientific Reports (Q1, IF 4.6)\n2025',
     'Risk factors focus; Title ngắn'),
    ('3', 'Association between screen time and depressive and anxiety symptoms '
     'among Chinese adolescents',
     'Anonymous (Trung Quốc)',
     'Frontiers in Psychiatry (Q1, IF 4.7)\n2025',
     'Single risk factor → outcomes'),
    ('4', 'The role of anxiety and trauma in predicting school avoidance among '
     'students: a structural equation modeling analysis',
     'Trung Quốc (cocấp lab)',
     'PubMed Central\n2025',
     '**SEM-focused title** + clear method'),
    ('5', 'Factors predicting the mathematics anxiety of adolescents: a '
     'structural equation modeling approach',
     'Bangladesh',
     'Frontiers in Psychiatry (Q1, IF 4.7)\n2024',
     '**"SEM approach" pattern** Q1 friendly'),
    ('6', 'Personality traits and psychological distress in Chinese '
     'adolescents: the mediating roles of anxiety and depression',
     'Trung Quốc',
     'Frontiers in Psychology (Q1, IF 3.7)\n2025',
     '**Mediation model title** standard format'),
    ('7', 'Social support and anxiety, a moderated mediating model',
     'Trung Quốc',
     'Scientific Reports (Q1, IF 4.6)\n2025',
     '**Moderated mediation** pattern - mới mẻ'),
    ('8', 'Gender differences and post-pandemic mental health impacts: a '
     'mediation study on Vietnamese adolescents',
     'Việt Nam (552 HS)',
     'Italian Journal of Medicine (Q2)\n2025',
     '**Title VN-specific** + mediation focus'),
    ('9', 'Trends and determinants of childhood anxiety disorders burden in '
     'Asia, 1990–2023',
     'Châu Á (multi-country)',
     'Journal of Affective Disorders (Q1, IF 6.6)\n2025',
     'Regional trend title pattern'),
    ('10', 'Effects of different interventions on anxiety disorders in '
     'children and adolescents: a systematic review and bayesian network '
     'meta-analysis',
     'Li L, Li Q, Wang J et al. (TQ)',
     'BMC Psychiatry (Q1, IF 4.4)\n2025',
     '**Intervention SR/MA** = high impact'),
    ('11', 'Mental health literacy as a moderator: association between '
     'psychological vulnerability and adolescent anxiety',
     'Trung Quốc (Guizhou, N=1591)',
     'PMC\n2025',
     'Moderator analysis title'),
    ('12', 'Association Between Comorbid Anxiety and Depression and Health '
     'Risk Behaviors Among Chinese Adolescents: Cross-Sectional Questionnaire '
     'Study',
     'Trung Quốc',
     'JMIR (Q1)\n2023',
     'Long descriptive title - cross-sectional pattern'),
]
make_title_table(q1_asia, [0.8, 6.5, 3.5, 3.5, 3.5])


H2('1.2 Châu Phi — Ethiopia, Nigeria, Rwanda, Sub-Saharan')

q1_africa = [
    ('13', 'Determinants of adolescents\' depression, anxiety, and somatic '
     'symptoms in Northwest Ethiopia: A non-recursive structural equation '
     'modeling',
     'Ethiopia (N=1,407)',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2023',
     '**"non-recursive SEM" title** đẹp + clear method'),
    ('14', 'Prevalence and correlates of anxiety and depressive symptoms among '
     'adolescents aged 10–19 years in six sub-Saharan African countries, '
     'China and India: A cross-sectional study',
     'Multi-country (N=9,849)',
     'PLOS Mental Health\n2025',
     'Cross-regional title — comprehensive'),
    ('15', 'The prevalence of mental health problems in sub-Saharan '
     'adolescents: A systematic review',
     'Sub-Saharan (review)',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2021',
     'Regional SR pattern'),
    ('16', 'COVID-19-related dysfunctional anxiety and associated factors '
     'among adolescents in Southwest Ethiopia: a cross-sectional study',
     'Ethiopia',
     'PMC\n2024',
     'Crisis-specific + region + cross-sectional'),
    ('17', 'Assessing anxiety symptom severity in Rwandese adolescents: '
     'cross-gender measurement invariance of GAD-7',
     'Rwanda (TQ-Africa lab)',
     'Frontiers in Psychiatry (Q1, IF 4.7)\n2024',
     '**Measurement invariance title — GẦN VỚI Q1 CỦA NHÓM**'),
    ('18', 'Generalized anxiety disorder screening using GAD-7 among '
     'in-school adolescents of Anambra State, Nigeria: a comparative study '
     'between urban and rural areas',
     'Nigeria',
     'PMC\n2023',
     'Urban-rural comparison, clear scope'),
]
make_title_table(q1_africa, [0.8, 6.5, 3.5, 3.5, 3.5])


d.add_page_break()


# ============================================================
H1('PHẦN 2 — Q3 / DESCRIPTIVE TIER: Cross-sectional Prevalence Studies')
P('Các bài Q3 / Q2 đặc trưng descriptive + prevalence + correlates, '
  'tiêu đề thường có cấu trúc "Prevalence and [associated factors / '
  'correlates / determinants]".', italic=True)

H2('2.1 Châu Á — Bangladesh, India, Pakistan, Sri Lanka')

q3_asia = [
    ('1', 'Prevalence of depression, anxiety and associated factors among '
     'school going adolescents in Bangladesh: Findings from a cross-sectional '
     'study',
     'Bangladesh (Dhaka)',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2021',
     '**Pattern điển hình Q3**: prevalence + AF + cross-sectional'),
    ('2', 'Anxiety among urban, semi-urban and rural school adolescents in '
     'Dhaka, Bangladesh: Investigating prevalence and associated factors',
     'Bangladesh',
     'PLOS ONE\n2022',
     'Urban-rural subgroup angle'),
    ('3', 'Prevalence of depression and anxiety among school going '
     'adolescents of Delhi: A cross-sectional study',
     'Ấn Độ (Delhi)',
     'PMC\n2025',
     'Single-city descriptive Q3 style'),
    ('4', 'Prevalence and socio-demographic correlates of depression and '
     'anxiety among late adolescents (15 to 21 years) in Mymensingh '
     'division, Bangladesh: A cross-sectional study',
     'Bangladesh (Mymensingh)',
     'PMC\n2025',
     'Regional + age range specific'),
    ('5', 'School-based intervention for anxiety using group cognitive '
     'behavior therapy in Pakistan: a feasibility randomized controlled trial',
     'Pakistan',
     'BMC\n2024',
     'Intervention angle title'),
    ('6', 'Assessment of mental health problems among adolescents in Sri '
     'Lanka: Findings from the cross-sectional Global School-based Health '
     'Survey',
     'Sri Lanka',
     'PMC\n2022',
     'Global health survey data'),
    ('7', 'Prevalence and associated factors of depressive and anxiety '
     'symptoms among Chinese secondary school students',
     'Trung Quốc',
     'PMC\n2023',
     'Standard prevalence title'),
    ('8', 'Prevalence and determinants of depression, anxiety, and stress '
     'among secondary school students',
     'Anonymous',
     'PLOS ONE (Q1/Q2, IF 3.7)\n2025',
     '**Q3 friendly title** simple structure'),
    ('9', 'Prevalence of internet addiction and anxiety, and factors '
     'associated with the high level of anxiety among adolescents in Hanoi, '
     'Vietnam during the COVID-19 pandemic',
     'Việt Nam (Hà Nội)',
     'PMC\n2023',
     '**VN-specific Hanoi title** - tham khảo gần với nhóm'),
    ('10', 'Depression, anxiety, and suicidal ideation among Vietnamese '
     'secondary school students and proposed solutions: a cross-sectional '
     'study',
     'Việt Nam',
     'PMC\n2014',
     'VN-secondary title classic pattern'),
    ('11', 'Mental health among ethnic minority adolescents in Vietnam and '
     'correlated factors: a cross-sectional study',
     'Việt Nam (DTTS)',
     'J Affective Disorders Reports (Q3)\n2024',
     'VN + subgroup focus'),
]
make_title_table(q3_asia, [0.8, 6.5, 3.5, 3.5, 3.5])


H2('2.2 Châu Phi — Ethiopia, Nigeria, etc.')

q3_africa = [
    ('12', 'Generalized anxiety disorder screening using GAD-7 among in-'
     'school adolescents of Anambra State, Nigeria: a comparative study '
     'between urban and rural areas',
     'Nigeria',
     'PMC\n2023',
     'Urban-rural pattern (lặp với #18 trên)'),
    ('13', 'Prevalence and Associated Factors of Depression, Anxiety, and '
     'Stress among High School Students in Northwest Ethiopia',
     'Ethiopia (Nakie et al. 2022)',
     'BMC Psychiatry (Q1, IF 4.4)\n2022',
     '**TÁC GIẢ ĐÃ CITED TRONG Q1+Q3 CỦA NHÓM** — '
     'mẫu Q3-style trên Q1 journal'),
    ('14', 'Screening for anxiety and its determinants among secondary school '
     'students during the COVID-19 era: a snapshot from Qatar in 2021',
     'Qatar',
     'PMC\n2022',
     'Crisis-era + region snapshot'),
    ('15', 'Anxiety related disorders in adolescents in the United Arab '
     'Emirates: a population based cross-sectional study',
     'UAE',
     'PMC\n2020',
     'Population-based descriptive'),
]
make_title_table(q3_africa, [0.8, 6.5, 3.5, 3.5, 3.5])


d.add_page_break()


# ============================================================
H1('PHẦN 3 — PHÂN TÍCH PATTERN')

H2('3.1 Cấu trúc Q1 titles thường dùng')
B('**SEM/Mediation + Population + Variable**: "[Variable] and [Outcome] in '
  '[Population]: a structural equation modeling analysis"')
B('Vd: "Determinants of adolescents\' depression... in Northwest Ethiopia: '
  'A non-recursive structural equation modeling"')
B('**Multi-group invariance**: "[Constructs] in [Population]: [statistical] '
  'invariance of [scale]"')
B('Vd: "Assessing anxiety symptom severity in Rwandese adolescents: '
  'cross-gender measurement invariance of GAD-7"')
B('**Risk-Protective Framework**: "[Risk/Protective] factors of [Outcome] '
  'in [Population]"')
B('Vd: "Risk factors of depressive and anxiety symptoms in Chinese '
  'adolescents"')
B('**Trend/Burden**: "[Time-period] trends of [outcome] among [population]"')
B('Vd: "Thirty-year trends of anxiety disorders among adolescents..."')

H2('3.2 Cấu trúc Q3 / Cross-sectional titles')
B('**"Prevalence + AF" Pattern (chuẩn nhất)**: "Prevalence and [associated '
  'factors / correlates / determinants] of [outcome] among [population] in '
  '[region]: a cross-sectional study"')
B('Vd: "Prevalence and determinants of depression, anxiety, and stress '
  'among secondary school students"')
B('**Subgroup comparison**: "[Outcome] among [subgroup1, subgroup2, subgroup3] '
  '[population] in [region]: Investigating prevalence and associated factors"')
B('Vd: "Anxiety among urban, semi-urban and rural school adolescents in '
  'Dhaka, Bangladesh"')
B('**Regional snapshot**: "[Outcome] among adolescents in [country/region]: '
  'a population-based cross-sectional study"')
B('Vd: "Anxiety related disorders in adolescents in the United Arab '
  'Emirates: a population-based cross-sectional study"')


d.add_page_break()


# ============================================================
H1('PHẦN 4 — ĐỀ XUẤT 10 TÊN BÀI CHO NHÓM CHỌN')

H2('4.1 Cho Q1 (BMC Psychiatry, Integrated SEM + Mixed-methods)')

q1_proposed = [
    ('A1', 'Integrated risk-protective structural equation model of anxiety '
     'disorder subtypes among Vietnamese lower secondary school students: A '
     'mixed-methods study',
     'Bám sát đề cương v3 hiện tại; rõ method'),
    ('A2', 'Risk and protective pathways to generalized, separation, and '
     'social anxiety subtypes among Vietnamese adolescents: A mixed-methods '
     'structural equation modeling study',
     'Nhấn 3 phân loại RLLA cụ thể trong title'),
    ('A3', 'Differential gender invariance across DSM-5 anxiety disorder '
     'subtypes among Vietnamese lower secondary students: An integrated SEM '
     'and qualitative study',
     'Nổi bật phát hiện CHÍNH (gender invariance SAD)'),
    ('A4', 'Multi-group structural equation modeling of risk and protective '
     'factors for anxiety disorder subtypes among Vietnamese adolescents: '
     'A mixed-methods cross-sectional study',
     '"Multi-group SEM" như Rwandese-GAD pattern'),
    ('A5', 'Beyond gender uniformity: A mixed-methods structural equation '
     'modeling analysis of anxiety disorder subtypes among Vietnamese lower '
     'secondary school students',
     'Title sáng tạo "Beyond gender uniformity" — eye-catching nhưng risky'),
]
make_title_table([(r[0], r[1], 'Hang Thi Cong et al. (VN)',
                   'BMC Psychiatry (Q1)', r[2]) for r in q1_proposed],
                 [0.8, 9.5, 2.5, 2.0, 3.0])


H2('4.2 Cho Q3 (PLOS ONE, Descriptive)')

q3_proposed = [
    ('B1', 'Manifestations and patterns of anxiety disorder subtypes among '
     'Vietnamese lower secondary school students: A descriptive cross-'
     'sectional study',
     'Bám sát đề cương v3 hiện tại'),
    ('B2', 'Prevalence and item-level patterns of generalized, separation, '
     'and social anxiety symptoms among Vietnamese lower secondary school '
     'students: A descriptive cross-sectional study',
     '**Pattern Q3 chuẩn** "Prevalence and..." như Bangladesh-PLOSONE'),
    ('B3', 'Item-level analysis of anxiety disorder subtypes and grade-level '
     'developmental trajectories among Vietnamese lower secondary school '
     'students: A descriptive normative study',
     'Nhấn "normative data" + "developmental"'),
    ('B4', 'Anxiety disorder subtypes among Vietnamese lower secondary '
     'students: A cross-sectional descriptive study of item-level patterns, '
     'gender, and grade trajectories',
     'Tổng hợp 3 RQ trong title — dài nhưng comprehensive'),
    ('B5', 'Generalized, separation, and social anxiety symptoms among '
     'Vietnamese adolescents in Hanoi: A descriptive item-level analysis for '
     'screening tool development',
     'Nhấn applied focus "screening tool development"'),
]
make_title_table([(r[0], r[1], 'Hang Thi Cong et al. (VN)',
                   'PLOS ONE', r[2]) for r in q3_proposed],
                 [0.8, 9.5, 2.5, 2.0, 3.0])


d.add_page_break()


# ============================================================
H1('PHẦN 5 — KHUYẾN NGHỊ EM')

H2('5.1 Em đề xuất top picks')

P('**Cho Q1**: A1 (giữ nguyên đề cương) HOẶC A4 (nhấn multi-group method)',
  italic=False)
P('  • Lý do A1: title cô đọng + nhấn 3 contributions (Integrated SEM + '
  'subtypes + mixed-methods)', indent=False)
P('  • Lý do A4: "Multi-group SEM" gần với Rwandese-GAD success story; có '
  'thể dễ accept hơn vì pattern đã được Q1 chấp nhận', indent=False)
P('  • A3 quá đặc thù (gender invariance là 1 phát hiện, không nên giành '
  'cả title)', indent=False)
P('  • A5 risky — "Beyond" trong title có thể bị reviewers cho là '
  'sensationalist', indent=False)

P('')
P('**Cho Q3**: B2 (Prevalence and item-level patterns) HOẶC B5 '
  '(screening tool development)', italic=False)
P('  • Lý do B2: đúng pattern PLOS ONE Q3 standard "Prevalence and...", '
  'matches Bangladesh 2021 successful pattern', indent=False)
P('  • Lý do B5: applied focus → reviewers PLOS ONE thích ứng dụng thực tiễn',
  indent=False)
P('  • B1 (current) OK nhưng "Manifestations and patterns" hơi vague — '
  'reviewer có thể hỏi "what manifestations specifically?"', indent=False)


H2('5.2 Câu hỏi thảo luận nhóm')
B('Q1 nên emphasize METHOD (Integrated SEM, mixed-methods) hay FINDING '
  '(gender invariance, differential pathways)?')
B('Có nên nêu rõ "Vietnam"/"Hanoi" trong title? — Reviewers thích vì rõ '
  'context; nhưng có thể giảm cross-cultural appeal')
B('"Lower secondary school" vs "middle school" vs "junior high" vs "early '
  'adolescents" — chọn term phù hợp tradeoff giữa specificity và universality')
B('"DSM-5 anxiety subtypes" vs "generalized, separation, social anxiety" — '
  'spell out có ưu/nhược gì?')
B('Title Q1+Q3 nên có format SIMILAR để dễ identify cùng dataset/program '
  '(vd: cùng kết thúc bằng "[...] among Vietnamese lower secondary school '
  'students")?')


H2('5.3 Em xin ý kiến thầy + đồng tác giả')

P('Em đã list 5 options cho Q1 + 5 options cho Q3 ở phần 4. Thầy + đồng '
  'tác giả có thể:', italic=True)
B('Chọn 1 option từ A1-A5 cho Q1 + 1 option từ B1-B5 cho Q3')
B('Hoặc đề xuất combination/sửa đổi (vd: "A1 nhưng đổi mixed-methods → '
  'qualitative-quantitative")')
B('Hoặc đưa ra hướng title hoàn toàn mới')


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
