"""
STEP 5: Cào mới 3 bài 2024-2026 từ web (chỉ abstract, chưa có PDF).
- VN007: Nguyen TL 2025 Italian J Med — Vietnam Gender + Post-pandemic
- QT059: Cai 2025 Frontiers Psychiatry — School Resilience SR/MA
- QT060: Bie 2024 Frontiers Psychiatry — Global GBD anxiety 10-24y 1990-2021
"""
import sys, io, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
TT_DIR = BASE / 'Tom-tat-tung-bai'

PAPERS = [
    {
        'id': 'VN007',
        'descriptor': 'NguyenTL_Gender_PostPandemic_ItalJMed_2025',
        'title_vn': 'Khác biệt giới tính và tác động của hậu đại dịch lên sức khỏe tâm thần: Nghiên cứu trung gian trên vị thành niên Việt Nam',
        'title_en': 'Gender differences and post-pandemic mental health impacts: a mediation study on Vietnamese adolescents',
        'authors': 'Nguyen TL, Storm V, Phung LT, Le CT, Le TTH, Nguyen TH, Vu ST, Tran TT, Vu MH, Nguyen VT',
        'affiliation': 'Liên kết Việt - quốc tế (cần verify trường cụ thể)',
        'journal': 'Italian Journal of Medicine (2025)',
        'year': 2025,
        'doi': '10.4081/itjm.2025.1888',
        'sample': '552 vị thành niên Việt Nam',
        'location': 'Việt Nam (chưa nêu thành phố cụ thể trong abstract)',
        'tool': 'RADS (Reynolds Adolescent Depression Scale) + SCARED (Screen for Child Anxiety-Related Emotional Disorders)',
        'design': 'Cắt ngang + phân tích trung gian (mediation analysis)',
        'outcomes': [
            'Tỷ lệ lo âu: 26,3 %',
            'Tỷ lệ trầm cảm: 13,8 %',
            'Lo âu chia tách (separation anxiety) phổ biến nhất: 33,0 %',
            'Lo âu xã hội (social phobia): 21,2 %',
            'Lo âu lan tỏa (GAD): 16,3 %',
            'Nữ giới có aOR = 5,7 cho trầm cảm (p < 0,01) — chênh GIỚI rất lớn',
            'Hậu đại dịch tăng trầm cảm: aOR = 4,3 (p < 0,001), tăng lo âu aOR = 3,8 (p < 0,01)',
            'KHÁM PHÁ TRUNG GIAN: thay đổi SKTT hậu đại dịch giải thích 60,8 % chênh lệch giới',
        ],
        'strengths': [
            'Mediation analysis — phương pháp hiện đại tách hiệu ứng trực tiếp / gián tiếp',
            'Sử dụng SCARED phân loại chi tiết LO ÂU theo subtype (separation, social, GAD)',
            'Phát hiện hậu đại dịch là cơ chế trung gian giới tính — có giá trị chính sách',
            'Bài VN trên tạp chí quốc tế peer-reviewed',
        ],
        'limitations': [
            'Cỡ mẫu vừa phải (n=552) so với V-NAMHS (5.996)',
            'Chưa rõ thành phố cụ thể trong abstract — cần verify từ full text',
            'Italian J Med — không phải tạp chí MH chuyên top — IF có thể thấp',
            'Cắt ngang → không xác định được nhân quả thật',
        ],
        'critique_short': 'Bài VIỆT NAM mới 2025 — phát hiện QUAN TRỌNG: hậu đại dịch giải thích 60,8 % chênh giới tính trong SKTT VTN. Phân loại lo âu theo SCARED subtype rất hữu ích (separation 33% > social 21,2 % > GAD 16,3 %). Bổ sung tốt cho VN029 Duong (TPHCM 2.631 HS) và VN014 Hoàng Trung Học (xu hướng COVID).',
        'future': 'Lặp lại với mẫu lớn hơn + nhiều thành phố. Theo dõi dọc xem mediation post-pandemic có suy giảm theo thời gian không.',
        'rating': 4,
        'notes_for_rag': 'VN 2025: lo âu 26,3%, trầm cảm 13,8%, separation anxiety nhất 33%, nữ aOR=5,7, post-pandemic mediates 60,8% gender gap.',
        'pdf': None,  # web only
        'pdf_folder': None,
        'source_url': 'https://www.italjmed.org/ijm/article/view/1888',
        'cao_date': '2026-04-24',
    },
    {
        'id': 'QT059',
        'descriptor': 'Cai_SchoolResilience_SR_MA_FrontPsych_2025',
        'title_vn': 'Can thiệp trường học để xây dựng tính kiên cường ở trẻ em và vị thành niên: Tổng quan hệ thống và phân tích tổng hợp các thử nghiệm RCT',
        'title_en': 'School-based interventions for resilience in children and adolescents: a systematic review and meta-analysis of randomized controlled trials',
        'authors': 'Cai C, Mei Z, Wang Z, Luo S',
        'affiliation': '(cần verify từ full text)',
        'journal': 'Frontiers in Psychiatry (2025)',
        'year': 2025,
        'doi': '10.3389/fpsyt.2025.1594658',
        'sample': '38 studies trong SR (21 trong MA), tổng 15.730 participants (8.869 can thiệp + 6.861 đối chứng)',
        'location': 'Toàn cầu (search nhiều databases)',
        'tool': 'Random-effects MA, SMD',
        'design': 'Systematic review + meta-analysis of RCTs',
        'outcomes': [
            'Hiệu quả TỔNG: SMD = 0,17 [95% CI: 0,06; 0,29], p < 0,01 — SMALL effect',
            'Phân nhóm theo loại can thiệp:',
            '  • Mindfulness-focused: SMD = 0,57 (LỚN nhất)',
            '  • Sport-focused: SMD = 0,41 (trung bình)',
            '  • Resilience-focused: SMD = 0,09 (rất nhỏ)',
            '  • Multi-component: SMD = 0,01 (gần như không hiệu quả)',
            '  • CBT approaches: SMD = -0,26 (NEGATIVE — cần verify từ full text!)',
            'Outcome chính là RESILIENCE (kiên cường), không phải lo âu/trầm cảm',
        ],
        'strengths': [
            '38 RCT — cỡ bằng chứng tốt',
            'Phân loại 5 nhóm can thiệp rõ ràng',
            'Frontiers Psychiatry Q1 IF',
            'Open Access ✓',
            'Phát hiện Mindfulness > Sport > Resilience > Multi > CBT — gợi ý ưu tiên đầu tư',
        ],
        'limitations': [
            'Outcome chính là RESILIENCE — không trực tiếp đo lo âu (cần đọc forest plot phụ về anxiety)',
            'CBT có SMD = -0,26 BẤT THƯỜNG — có thể là artifact của ít RCT, hoặc CBT thiếu component build resilience',
            'SMD tổng = 0,17 chỉ là SMALL — cần cost-effectiveness analysis',
            'Không tách theo độ tuổi (TE vs VTN) trong abstract',
        ],
        'critique_short': 'SR/MA QUAN TRỌNG về RESILIENCE — phát hiện gây ngạc nhiên: Mindfulness (SMD=0,57) tốt hơn CBT (SMD=-0,26). Nhưng CBT âm có thể do: (1) ít RCT đủ chất lượng đo resilience; (2) CBT design tập trung giảm triệu chứng, không build resilience tích cực. Cần đọc full text để hiểu CBT bị âm là do bias hay thật. Cho VN: gợi ý xây dựng can thiệp Mindfulness trường học hơn là CBT thuần.',
        'future': 'Đọc full text PDF (open access) để xác minh CBT âm. Nghiên cứu tương tự cho VN với HS THCS/THPT, dùng thang đo resilience VN-validated.',
        'rating': 4,
        'notes_for_rag': '38 RCT resilience school SR/MA. Mindfulness SMD 0,57 > Sport 0,41 > Resilience 0,09 > Multi 0,01 > CBT -0,26 (bất thường). Tổng 0,17 small.',
        'pdf': None,
        'pdf_folder': None,
        'source_url': 'https://pmc.ncbi.nlm.nih.gov/articles/PMC12127306/',
        'cao_date': '2026-04-24',
    },
    {
        'id': 'QT060',
        'descriptor': 'Bie_GlobalAnxiety_GBD_10-24y_1990_2021_FrontPsych_2024',
        'title_vn': 'Gánh nặng toàn cầu của rối loạn lo âu ở vị thành niên và thanh niên 10-24 tuổi 1990-2021: Xu hướng, yếu tố nguy cơ, và tác động chênh lệch kinh tế-xã hội + COVID-19',
        'title_en': 'Rising global burden of anxiety disorders among adolescents and young adults: trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021',
        'authors': 'Bie F, Yan X, Xing J, Wang L, Xu Y, Wang G, Wang Q, Guo J, Qiao J, Rao Z',
        'affiliation': 'Trung Quốc (theo names — cần verify)',
        'journal': 'Frontiers in Psychiatry (2024)',
        'year': 2024,
        'doi': '10.3389/fpsyt.2024.1489427',
        'sample': 'GBD 2021 — 204 quốc gia, độ tuổi 10-24',
        'location': 'Toàn cầu — phân tầng theo SDI và khu vực',
        'tool': 'GBD 2021 dataset, Joinpoint regression, age-standardized rates',
        'design': 'Phân tích thứ cấp dữ liệu GBD 2021',
        'outcomes': [
            'Tăng 52 % số ca lo âu mới ở 10-24 tuổi toàn cầu (1990 → 2021)',
            'Tỷ lệ mới mắc: 708 → 883 / 100.000 dân',
            'Nhóm 20-24 tuổi tăng nhiều nhất: +28,33 %',
            'Nhóm 10-14 tuổi DUY TRÌ tỷ lệ cao nhất xuyên suốt',
            'Nữ > nam ở mọi nhóm tuổi (consistent)',
            'Khu vực: Latin Mỹ Tropical cao nhất; Đông Á giảm nhiều nhất',
            'Nam Á dẫn đầu số ca tuyệt đối: 3,5 triệu (Ấn Độ chủ yếu)',
            'COVID-19 tăng tốc rõ rệt 2019-2021',
            'Yếu tố nguy cơ chính: BULLYING victimization (gánh nặng disability lớn nhất)',
        ],
        'strengths': [
            'GBD 2021 — dữ liệu cập nhật mới nhất',
            'Phân tầng nhiều chiều: tuổi, giới, khu vực, SDI, COVID',
            '32 năm theo dõi 1990-2021',
            'Frontiers Psychiatry Q1 Open Access ✓',
            'Ngoại suy mạnh cho VN (LMIC Đông Nam Á)',
        ],
        'limitations': [
            'GBD dùng modeling (DisMod-MR) — không phải đo trực tiếp',
            'Tin cậy thấp ở các nước có dữ liệu sơ sài (nhiều LMIC)',
            'Lo âu trong GBD là "anxiety disorders" tổng hợp — không tách subtype',
            'Bullying là proxy yếu tố nguy cơ — không factor analysis chi tiết',
        ],
        'critique_short': 'Bài QUAN TRỌNG để có baseline GLOBAL trends 10-24 tuổi 1990-2021. Xu hướng +52 % toàn cầu, nhưng tách theo khu vực: Đông Á GIẢM (Hàn Quốc + Nhật) trái với Tây phương tăng. VN ít data direct nhưng Nam Á + ASEAN có tăng. Cho thầy: dùng GBD trends làm BACKDROP cho luận điểm "VN cần khảo sát quốc gia lặp lại để có xu hướng riêng".',
        'future': 'So sánh trends VN với Hàn Quốc (giảm) và Tây phương (tăng) để xác định mô hình VN. Phân tích bullying VN qua các NC hiện có (VN015, VN022 UNICEF).',
        'rating': 5,
        'notes_for_rag': 'GBD 2021 anxiety 10-24y: +52% global 1990-2021, từ 708 → 883 per 100k. Nữ > nam, Latin Tropical cao nhất, Đông Á giảm, Nam Á số tuyệt đối lớn nhất, bullying yếu tố chính.',
        'pdf': None,
        'pdf_folder': None,
        'source_url': 'https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1489427/full',
        'cao_date': '2026-04-24',
    },
]

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def build_summary_docx(paper, out_path):
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    t = d.add_heading(f"{paper['id']} — Tóm tắt bài nghiên cứu (cào từ web)", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)

    # Note: no PDF
    note = d.add_table(rows=1, cols=1)
    note.style = 'Table Grid'
    nc = note.rows[0].cells[0]
    shade_cell(nc, 'FFF2CC')
    p = nc.paragraphs[0]
    r = p.add_run('Lưu ý: ')
    r.bold = True
    r.font.color.rgb = RGBColor(191, 97, 14)
    p.add_run(f"Bài này được CÀO từ web (abstract + key findings) ngày {paper['cao_date']}. CHƯA có PDF gốc trong thư viện. ")
    p.add_run(f"Cần tải full PDF từ {paper['source_url']} (open access) để verify số liệu chi tiết.")
    d.add_paragraph()

    meta_tbl = d.add_table(rows=8, cols=2)
    meta_tbl.style = 'Table Grid'
    meta_rows = [
        ('Tên (VN)', paper['title_vn']),
        ('Tên (EN)', paper['title_en']),
        ('Tác giả', paper['authors']),
        ('Đơn vị', paper['affiliation']),
        ('Nơi công bố', paper['journal']),
        ('Năm', str(paper['year'])),
        ('DOI', paper['doi']),
        ('Mẫu + công cụ', f"{paper['sample']} | Công cụ: {paper['tool']}"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        cell0 = meta_tbl.rows[i].cells[0]
        shade_cell(cell0, 'D9E1F2')
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        meta_tbl.rows[i].cells[1].text = v
    d.add_paragraph()

    for sec, items, color in [
        ('Phương pháp', [paper['design'], f"Mẫu: {paper['sample']}", f"Công cụ: {paper['tool']}"], (31, 73, 125)),
        ('Kết quả nghiên cứu', paper['outcomes'], (31, 73, 125)),
        ('Điểm mạnh', paper['strengths'], (54, 95, 44)),
        ('Hạn chế', paper['limitations'], (192, 80, 77)),
    ]:
        h = d.add_heading(sec, level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor(*color)
        for it in items:
            d.add_paragraph('• ' + it)

    h = d.add_heading('Phản biện ngắn', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(192, 0, 0)
    p = d.add_paragraph()
    r = p.add_run(paper['critique_short'])
    r.font.color.rgb = RGBColor(192, 0, 0)

    h = d.add_heading('Hướng nghiên cứu tiếp theo', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)
    d.add_paragraph(paper['future'])

    stars = '⭐' * paper['rating']
    h = d.add_heading(f"Đánh giá chất lượng: {stars} ({paper['rating']}/5)", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(191, 97, 14)

    if paper.get('notes_for_rag'):
        d.add_paragraph()
        p = d.add_paragraph()
        r = p.add_run('Key fact cho RAG: ')
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(paper['notes_for_rag']).font.size = Pt(10)

    d.save(out_path)

# Execute
log = []
with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

for paper in PAPERS:
    tt_name = f"{paper['id']}_{paper['descriptor']}.docx"
    tt_path = TT_DIR / tt_name
    build_summary_docx(paper, tt_path)
    log.append(f"SUMMARY {paper['id']}: {tt_name} ({tt_path.stat().st_size//1024} KB)")

    idx[paper['id']] = {
        'id': paper['id'],
        'descriptor': paper['descriptor'],
        'summary': tt_name,
        'translation': None,
        'pdf': None,
        'pdf_folder': None,
        'doi': paper.get('doi'),
        'source_url': paper.get('source_url'),
        'cao_date': paper.get('cao_date'),
        'scope': 'web_caugh_no_pdf',
    }
    log.append(f"INDEX {paper['id']}: added (no PDF, source_url logged)")

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

n_vn = sum(1 for k in idx if k.startswith('VN'))
n_qt = sum(1 for k in idx if k.startswith('QT'))
log.append(f'\nTotal canonical: {len(idx)} ({n_vn} VN + {n_qt} QT)')
for l in log:
    print(l)
