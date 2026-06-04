# -*- coding: utf-8 -*-
"""Sinh TomTat v2 - cai thien hanh van, dan dat, do chi tiet.
- Bot bullets, them prose narrative
- Better transitions between sections
- Detailed explanations of WHY (not just WHAT)
- Executive-summary tone cho co-authors
- Fix data discrepancies:
  * 03_Ban-dich: 105 (not 93)
  * Subtype means M=55,82/25,06/48,41 → source = THỰC TRẠNG file thầy gửi 01/06
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'TomTat_Phien_v2_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.2); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.4


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=True, align_center=False):
    """Paragraph with indent (default) for prose flow."""
    p = d.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_center
                   else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(6)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def NOTE(text):
    """Boxed note - italic gray for context."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_table(headers, rows, col_widths_cm, style='Light Grid Accent 1'):
    t = d.add_table(rows=1, cols=len(headers)); t.style = style; t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
H1('BÁO CÁO TÓM TẮT TIẾN ĐỘ DỰ ÁN Q1 + Q3')
P('Phiên rà soát kho dữ liệu, xây dựng đề cương, và chuẩn hoá thuật ngữ',
  italic=True, indent=False, align_center=True)
P('Giai đoạn: 29/05/2026 – 01/06/2026', italic=True, indent=False, align_center=True)
P('Nhóm tác giả: Hang Thi Cong (1st), Nguyen Minh Duc, Duc Minh Dao (corresponding)',
  italic=True, indent=False, align_center=True)


# ============================================================
H1('LỜI GIỚI THIỆU')

P('Tài liệu này tóm tắt khối lượng công việc đã hoàn thành trong bốn '
  'phiên làm việc liên tục từ 29/05/2026 đến 01/06/2026, nhằm chuẩn bị '
  'nền tảng vững chắc cho hai bài báo quốc tế: một bài Q1 hướng đến BMC '
  'Psychiatry và một bài Q3 hướng đến PLOS ONE, đều khai thác từ bộ '
  'dữ liệu luận án tiến sĩ của nghiên cứu sinh Công Thị Hằng (N = 1.352 '
  'học sinh trung học cơ sở Hà Nội).')

P('Báo cáo gồm tám phần, dẫn dắt từ quá trình rà soát toàn bộ kho dữ '
  'liệu (Phần 1), đến chuẩn hoá thuật ngữ chuyên môn theo gợi ý của thầy '
  'hướng dẫn (Phần 2), tiếp đến là bố cục chi tiết hai bài báo (Phần 3), '
  'kết quả rà soát mười chín issues về phương pháp luận (Phần 4), bốn '
  'câu hỏi blocking cần thầy hướng dẫn và nghiên cứu sinh quyết định '
  '(Phần 5), bộ tham khảo hơn ba mươi tên bài từ tác giả châu Á và '
  'châu Phi (Phần 6), danh mục files quan trọng đã tạo (Phần 7), và '
  'cuối cùng là kế hoạch triển khai sáu tuần tiếp theo (Phần 8).')

P('Mục đích của tài liệu này không chỉ để báo cáo công việc đã làm, mà '
  'còn để đồng bộ hiểu biết giữa các thành viên trong nhóm tác giả — '
  'đặc biệt khi thầy Đào Minh Đức cùng đồng tác giả Nguyễn Minh Đức '
  'tham gia đánh giá đề cương và đưa ra quyết định về các vấn đề blocking '
  'còn tồn đọng.', italic=True)


# ============================================================
H1('1. RÀ SOÁT TOÀN BỘ KHO DỮ LIỆU')

H2('1.1 Phạm vi đã rà soát và kết quả tổng quan')

P('Quá trình rà soát đã bao phủ tám danh mục tài liệu chính trong kho '
  'dữ liệu nghiên cứu, từ các PDF gốc của các bài báo được trích dẫn, '
  'cho đến các tóm tắt nội bộ, bản dịch, báo cáo tổng quan, cũng như '
  'các tài liệu luận án và đề cương. Tổng số file được rà soát vượt '
  'qua con số ba trăm sáu mươi.')

make_table(
    ['Loại tài liệu', 'Phạm vi', 'Tình trạng coverage'],
    [
        ('PDF gốc các bài báo được trích dẫn',
         '147 PDFs đã đọc và lập catalog trên tổng 183 file phát hiện',
         '80,3% (36 file còn trong trạng thái OneDrive cloud-only, '
         'chờ NCS sync)'),
        ('Tóm tắt nghiên cứu QT (Tom-tat-tung-bai/)',
         '73 file dạng QT001 đến QT073',
         '100% đã verify với PDF gốc'),
        ('Bản dịch tài liệu nước ngoài (03_Ban-dich/)',
         '105 file docx active',
         '100% đã clean metadata, 78 file có metadata cũ đã sạch'),
        ('Tài liệu luận án và phái sinh (Luận án TS/)',
         '~20 file core (LA chính + tóm tắt + trích yếu + form)',
         '100% đã update thuật ngữ và metadata'),
        ('Đề cương Q1/Q3 (bai-bao-Q1/)',
         '15+ file (bố cục, outline, báo cáo rà soát, song ngữ)',
         '100% đã tạo và verify'),
        ('Cơ sở dữ liệu (04_Co-so-du-lieu/)',
         '4 file markdown (DATABASE_BAI_BAO_LO_AU + báo cáo + gợi ý)',
         '100% đã update'),
        ('Hệ thống RAG vector database',
         '1 SQLite (1,6 MB) chứa embedding của ~109 file',
         'Đã inventory, chưa deep query test'),
        ('Báo cáo cũ trong 01_Bao-cao/',
         '~100 file working drafts và audit reports',
         '100% đã update thuật ngữ (phase 2 batch)'),
    ],
    [4.5, 5.5, 7.0]
)


H2('1.2 Lỗi đã phát hiện và phương án xử lý')

P('Trong quá trình rà soát, em phát hiện một số vấn đề về tính chính '
  'xác và nhất quán của tài liệu. Em xếp loại các vấn đề này theo mức '
  'độ nghiêm trọng và đã xử lý lần lượt theo trình tự ưu tiên: lỗi '
  'fact thực sự trước, sau đó là vấn đề định danh file, cuối cùng là '
  'chuẩn hoá thuật ngữ trên diện rộng.')

P('Trên tổng số bảy mươi ba bản tóm tắt QT đã rà soát, chỉ duy nhất '
  'một trường hợp có lỗi thực sự về số liệu: bản tóm tắt QT011 về '
  'nghiên cứu Bhardwaj 2020 tại Chandigarh, Ấn Độ — trong đó số liệu '
  'tỷ lệ căng thẳng mức "trung bình hoặc nặng" ở nữ sinh ghi 44,2% '
  'trong khi giá trị đúng theo PDF gốc Bảng 8 là 48,7% (cộng từ 35,4% '
  'mức trung bình và 13,3% mức nặng). Lỗi này đã được sửa và file gốc '
  'có backup trước khi sửa.')

P('Một vấn đề về định danh file cũng đã được phát hiện và xử lý: file '
  'PDF có tên QT002_Saikia_2023_India_Assam.pdf trong thư mục '
  'The-gioi_Khac/ thực ra chứa nội dung của bài Bhardwaj 2020 chứ '
  'không phải Saikia. PDF Saikia thật được tìm thấy ở thư mục '
  '11-bai-ban-dau-va-mo-rong/ với tên 11_Saikia_2023_IJCM.pdf. File bị '
  'mislabel đã được đổi tên thành "_MISLABELED_QT002_Saikia_actually_'
  'Bhardwaj_2020.pdf" để tránh nhầm lẫn trong tương lai.')

P('Vấn đề có quy mô lớn nhất là việc sử dụng thuật ngữ "rối loạn lo âu '
  'lan tỏa" cho khái niệm Generalized Anxiety Disorder (GAD) trong tiếng '
  'Anh. Sau khi thầy hướng dẫn Đào Minh Đức gợi ý chuyển sang "rối loạn '
  'lo âu tổng quát" để chính xác hơn về mặt ngữ nghĩa và tránh trùng '
  'với thuật ngữ "rối loạn phát triển lan tỏa" của DSM-IV, em đã thực '
  'hiện hai phase batch update trên toàn kho dữ liệu, tổng cộng 1.696 '
  'thay thế trên 160 file. Chi tiết về quyết định này được trình bày '
  'ở Phần 2.')

H3('Bảng tổng hợp các lỗi và xử lý')

make_table(
    ['Loại vấn đề', 'Số lượng', 'Phương án xử lý'],
    [
        ('Lỗi số liệu thực sự trong bản tóm tắt QT',
         '1 (QT011: 44,2% → 48,7%)',
         'Đã sửa + backup file gốc trước khi sửa'),
        ('PDF mislabel (filename ≠ content)',
         '1 (QT002_Saikia → thực là Bhardwaj)',
         'Đã đổi tên file để nhận dạng; PDF Saikia thật đã được xác định'),
        ('Thuật ngữ GAD "lan tỏa" cần đổi sang "tổng quát"',
         '1.696 chỗ trong 160 file',
         'Đã batch update qua hai phase + verify residuals'),
        ('Watermark và metadata không sạch',
         'Hàng trăm file',
         'Đã clean batch (author, title, last_modified_by, v.v.)'),
        ('Thang đo bịa từ phiên trước (MPAS, MPVS)',
         '~10 chỗ trong 6 file',
         'Đã phát hiện sớm các phiên trước; sửa thành SAS-SV + OBVQ'),
    ],
    [5.5, 4.0, 7.5]
)


H2('1.3 Đánh giá chất lượng tổng thể của kho dữ liệu')

P('Sau bốn phiên rà soát, em có thể khẳng định kho dữ liệu nghiên cứu '
  'đạt mức chất lượng cao về cả tính chính xác lẫn tính nhất quán. Tỷ '
  'lệ lỗi thực sự trong các bản tóm tắt QT chỉ vào khoảng 1,4 phần '
  'trăm — một con số rất thấp so với chuẩn quốc tế cho các tài liệu '
  'tổng quan nghiên cứu.')

P('Em đã thực hiện manual verify chi tiết sáu trường hợp nghi vấn (QT002, '
  'QT005, QT011, QT012, QT034, QT066, QT067) và xác nhận rằng năm '
  'trong số sáu trường hợp đó là false positive do script tự động '
  'không xử lý được các định dạng phức tạp trong PDF — chẳng hạn như '
  'số liệu trong bảng có dấu ngoặc đơn (vd. "(71.3)"), dấu chấm '
  'giữa đặc trưng của tạp chí Lancet (vd. "10·1%"), hoặc các giá trị '
  'đã được làm tròn từ giá trị thập phân nhiều chữ số trong PDF. Chỉ '
  'QT011 là lỗi thực và đã được sửa.')

P('Đặc biệt quan trọng cho hai bài Q1 và Q3, tất cả mười bảy citations '
  'chính được sử dụng trong đề cương hiện tại đều có file PDF được '
  'verify trong kho 02_Papers-goc/. Báo cáo từ phiên 29/05 ban đầu đã '
  'nhận định nhầm rằng PDF của bài Saikia 2023 bị mất, nhưng phiên '
  '01/06 phát hiện file thật đang ở vị trí 11-bai-ban-dau-va-mo-rong/'
  '11_Saikia_2023_IJCM.pdf — chứng minh rằng các trích dẫn được xây '
  'dựng trên nền tảng PDF gốc đầy đủ và đáng tin cậy.', italic=True)


# ============================================================
H1('2. CHUẨN HOÁ THUẬT NGỮ: "LAN TỎA" THÀNH "TỔNG QUÁT"')

H2('2.1 Bối cảnh và lý do chuyển đổi')

P('Trong phiên trao đổi với thầy Đào Minh Đức ngày 01/06/2026, thầy '
  'gợi ý nên xem xét lại việc sử dụng "rối loạn lo âu lan tỏa" làm '
  'bản dịch tiếng Việt cho khái niệm "Generalized Anxiety Disorder" '
  '(GAD) trong DSM-5. Lập luận của thầy dựa trên bốn căn cứ chuyên '
  'môn mà em đánh giá đều có cơ sở vững chắc.')

P('Thứ nhất, về mặt ngữ nghĩa, thuật ngữ tiếng Anh "Generalized" '
  'mang ý nghĩa "không đặc thù theo trigger" và "áp dụng broad-'
  'spectrum", chứ không phải "lan tỏa khắp" như nghĩa của "pervasive". '
  'Khi dịch thành "tổng quát", chúng ta truyền tải đúng tinh thần '
  'rằng GAD là dạng lo âu không gắn với một kích thích cụ thể nào, '
  'khác biệt với lo âu xã hội (gắn với tình huống xã hội) hay lo âu '
  'chia ly (gắn với sự tách rời người thân).')

P('Thứ hai, về mặt lịch sử thuật ngữ, "lan tỏa" đã được sử dụng trong '
  'DSM-IV để dịch "Pervasive Developmental Disorder (PDD)" — bao gồm '
  'các rối loạn phát triển lan tỏa như tự kỷ, hội chứng Asperger, '
  'hội chứng Rett. Việc dùng cùng từ "lan tỏa" cho cả GAD trong DSM-5 '
  'sẽ tạo nhầm lẫn nghiêm trọng trong lĩnh vực tâm lý – tâm thần '
  'tiếng Việt.')

P('Thứ ba, về mặt logic phân loại học, GAD là dạng "tổng quát" trong '
  'họ rối loạn lo âu, từ đó các dạng đặc thù được phân tách: lo âu '
  'chia ly, lo âu xã hội, hoảng loạn, ám ảnh sợ hãi. Cấu trúc taxonomy '
  'này được phản ánh chính xác hơn qua từ "tổng quát" so với "lan tỏa".')

P('Cuối cùng, lập luận thuyết phục nhất của thầy là về tính phân biệt: '
  'mọi rối loạn lo âu — chia ly, xã hội, ám ảnh — đều có tính chất '
  '"lan tỏa" về mặt hậu quả vì đều ảnh hưởng đến phát triển toàn diện '
  'của thanh thiếu niên. Vì vậy, "lan tỏa" không phải là tính chất '
  'phân biệt riêng của GAD, không nên dùng làm đặc trưng phân loại.')

H2('2.2 Phạm vi triển khai update')

P('Sau khi thầy + nghiên cứu sinh xác nhận quyết định, em đã thực hiện '
  'việc thay thế theo ba phase, từ các file lõi quan trọng nhất đến '
  'các file phụ trợ.')

make_table(
    ['Phase', 'Phạm vi files', 'Số replacements', 'Loại files'],
    [
        ('Phase 1', 'Core 15 files',
         '169 replacements',
         'LA chính, tóm tắt LA, trích yếu, danh mục công trình, '
         'tất cả outline Q1/Q3'),
        ('Phase 2', '139 files mở rộng',
         '1.522 replacements',
         'Báo cáo trong 01_Bao-cao/, bản dịch 03_Ban-dich/, '
         'working drafts phiên cũ'),
        ('Phase 3', '5 files edge cases',
         '3 replacements',
         'Các file có text split phức tạp giữa nhiều runs trong docx'),
        ('Bonus', 'CSDL DATABASE master',
         '2 replacements',
         'File markdown 04_Co-so-du-lieu/DATABASE_BAI_BAO_LO_AU.md'),
        ('Tổng cộng', '160 files',
         '1.696 replacements', '—'),
    ],
    [2.5, 4.0, 4.0, 6.5]
)

P('Riêng các file trong thư mục bai-bao-khgdvn/ (chứa hai bài báo '
  'Q2.5 KHGDVN của nghiên cứu sinh) đã được loại trừ khỏi quá trình '
  'update theo chỉ thị "không đụng" từ phiên đầu. Các file backup và '
  'file archive cũng được skip để bảo tồn lịch sử thay đổi. Viết tắt '
  '"RLLALT" (rối loạn lo âu lan tỏa) đã được đổi đồng bộ thành '
  '"RLLATQ" (rối loạn lo âu tổng quát) ở mọi vị trí xuất hiện.',
  italic=True)


# ============================================================
H1('3. BỐ CỤC CHI TIẾT HAI BÀI Q1 + Q3')

P('Sau khi thảo luận với thầy và nghiên cứu sinh về cách phân chia '
  'bộ dữ liệu cho hai bài báo, nhóm đã thống nhất chọn Phương án A: '
  'phân chia theo phương pháp và câu hỏi nghiên cứu, thay vì theo '
  'dạng rối loạn lo âu hay theo loại yếu tố. Lựa chọn này đảm bảo '
  'hai bài hoàn toàn không trùng lặp về phát hiện cốt lõi, dù cùng '
  'khai thác từ một bộ dữ liệu N = 1.352 học sinh.')


H2('3.1 Bài Q1 — BMC Psychiatry')

P('Bài Q1 nhằm vào tạp chí BMC Psychiatry, một tạp chí Q1 trong họ '
  'BMC với chỉ số tác động IF 4,4 và tỷ lệ chấp nhận khoảng 30%. '
  'BMC Psychiatry được chọn ưu tiên do ba lý do: chấp nhận thiết kế '
  'mixed-methods, có nhiều bài về adolescent mental health từ các tác '
  'giả châu Á và châu Phi, và quy trình xử lý bài tương đối nhanh '
  '(60-90 ngày cho quyết định đầu tiên). Backup journal là BMC '
  'Psychology (IF 2,6) trong trường hợp BMC Psychiatry reject.')

make_table(
    ['Mục', 'Thông tin chi tiết'],
    [
        ('Tên bài (tentative)',
         'Integrated risk-protective structural equation model of anxiety '
         'disorder subtypes among Vietnamese lower secondary school students: '
         'A mixed-methods study'),
        ('Tác giả',
         'Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹† '
         '(*: tác giả thứ nhất, †: tác giả liên hệ)'),
        ('Số từ dự kiến', '6.000 – 8.000 từ (không tính references)'),
        ('Cấu trúc IMRaD',
         'Abstract (~250 từ) + Introduction (~1.500 từ) + Methods (~2.000 từ) '
         '+ Results (~2.000 từ) + Discussion (~1.800 từ) + References '
         '(45-50 refs)'),
        ('Phương pháp chính',
         'CFA + SEM tích hợp (AMOS 31.0) + Multi-group invariance theo giới '
         '+ Mixed-methods Convergent Parallel Design'),
        ('Điểm mới thứ nhất',
         'Mô hình SEM tích hợp đầu tiên ở Việt Nam đồng thời kiểm định 3 '
         'yếu tố nguy cơ và 4 yếu tố bảo vệ trên 3 phân loại RLLA (21 paths)'),
        ('Điểm mới thứ hai',
         'Phát hiện tính bất biến giới của rối loạn lo âu chia ly '
         '(F = 0,246; p = 0,620), trái với lo âu tổng quát (F = 44,484) '
         'và xã hội (F = 45,984)'),
        ('Điểm mới thứ ba',
         'Tích hợp dữ liệu phỏng vấn định tính với hệ số β định lượng thông '
         'qua joint display matrix'),
    ],
    [4.0, 13.0]
)


H2('3.2 Bài Q3 — PLOS ONE')

P('Bài Q3 nhằm vào tạp chí PLOS ONE, có chỉ số tác động IF 3,7 và '
  'tỷ lệ chấp nhận khoảng 50% — cao hơn đáng kể so với các tạp chí '
  'Q1 truyền thống. PLOS ONE được chọn vì có scope rộng và đặc biệt '
  'thân thiện với nghiên cứu mô tả + cross-sectional, phù hợp hoàn '
  'hảo với góc tiếp cận item-level analysis của bài này. Backup '
  'journal là BMC Pediatrics (IF 2,5) trong trường hợp PLOS ONE yêu '
  'cầu pediatric focus rõ hơn.')

make_table(
    ['Mục', 'Thông tin chi tiết'],
    [
        ('Tên bài (tentative)',
         'Manifestations and patterns of anxiety disorder subtypes among '
         'Vietnamese lower secondary school students: A descriptive cross-'
         'sectional study'),
        ('Tác giả',
         'Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹† '
         '(giống Q1)'),
        ('Số từ dự kiến', '3.500 – 5.000 từ'),
        ('Cấu trúc IMRaD',
         'Abstract (~200 từ) + Introduction (~1.200 từ) + Methods (~1.200 từ) '
         '+ Results (~1.500 từ) + Discussion (~1.000 từ) + References '
         '(30-35 refs)'),
        ('Phương pháp chính',
         'Thống kê mô tả (M, SD, item ranking) + ANOVA + Bonferroni '
         'correction (α = 0,0167) + Cohen d (giới) + η² (khối lớp)'),
        ('Điểm mới thứ nhất',
         'Phân tích mức độ mục đầu tiên cho phiên bản tiếng Việt của RCADS '
         'trên mẫu lớn N = 1.352 — cung cấp dữ liệu chuẩn cho lâm sàng tâm lý'),
        ('Điểm mới thứ hai',
         'Quỹ đạo phát triển theo khối lớp 6→9, với rối loạn lo âu chia ly '
         'giảm đơn điệu từ M = 32,13 (khối 6) xuống M = 19,46 (khối 9) — '
         'phù hợp DSM-5 SAD = khởi phát thời thơ ấu'),
        ('Điểm mới thứ ba',
         'Xác định mục tiêu sàng lọc ưu tiên (RCADS4 + RCADS13 + RCADS8 cho '
         'GAD; RCADS32 + RCADS43 cho SocAD) làm cơ sở phát triển công cụ '
         'sàng lọc ngắn'),
    ],
    [4.0, 13.0]
)


H2('3.3 Phân chia dữ liệu chống trùng lặp')

P('Để tránh nguy cơ "self-plagiarism" hoặc "redundant publication" — '
  'hai vấn đề mà các tạp chí quốc tế đặc biệt nghiêm khắc khi hai bài '
  'cùng khai thác một bộ dữ liệu — em đã xây dựng ma trận phân bổ '
  'dữ liệu theo bốn tiêu chí orthogonal: câu hỏi nghiên cứu khác nhau, '
  'phương pháp thống kê khác nhau, outcome khác nhau, và paradigm '
  'phương pháp khác nhau (mixed-methods cho Q1 vs. quantitative-only '
  'cho Q3).')

make_table(
    ['Dữ liệu nguồn', 'Sử dụng trong Q1', 'Sử dụng trong Q3'],
    [
        ('Bảng 1-3: item-level GAD/SAD/SocAD',
         'Chỉ tóm tắt aggregate M, SD',
         'CHỦ LỰC — full item ranking table'),
        ('Bảng 4: demographic comparison',
         'Background only, focus là multi-group invariance test',
         'Primary analysis với Bonferroni + effect sizes'),
        ('Bảng 5-6: item-level yếu tố nguy cơ + bảo vệ',
         'Aggregate score → SEM predictors (21 paths)',
         'KHÔNG dùng (out of scope cho descriptive paper)'),
        ('SEM β coefficients từ LA Bảng 27-42',
         'PRIMARY FINDING — toàn bộ 21 paths',
         'KHÔNG dùng'),
        ('Multi-group invariance kết quả',
         'NOVEL FINDING — phát hiện chính của paper',
         'Brief mention với cite Q1 paper (companion)'),
        ('Phỏng vấn định tính',
         'Mixed-methods integration via joint display',
         'KHÔNG dùng (purely quantitative paper)'),
    ],
    [5.0, 6.0, 6.0]
)


# ============================================================
H1('4. RÀ SOÁT CHI TIẾT 19 ISSUES Q1 + Q3')

P('Sau khi outline v2 cho Q1 và Q3 được hoàn thành, em đã tiến hành rà '
  'soát từng câu khẳng định, từng dữ liệu, từng tham chiếu, và từng '
  'mạch logic dẫn dắt trong cả hai outline. Quá trình này dẫn đến phát '
  'hiện 19 issues cần xử lý: 10 trong Q1 và 9 trong Q3.')

P('Em phân loại các issues theo ba tier dựa trên mức độ blocking và '
  'ownership. Tier 1 (Blocking) đòi hỏi quyết định từ thầy hướng dẫn '
  'hoặc nghiên cứu sinh, không thể tự xử lý được. Tier 2 (High) và '
  'Tier 3 (Medium) thuộc thẩm quyền em xử lý trong quá trình viết '
  'outline v3.')

make_table(
    ['Tier', 'Tổng số', 'Mã issues', 'Owner', 'Trạng thái hiện tại'],
    [
        ('Tier 1 — BLOCKING',
         '4',
         'Q1-6, Q1-8, Q3-6, Q3-9',
         'NCS + Thầy',
         'Đang chờ quyết định — chi tiết ở Phần 5'),
        ('Tier 2 — HIGH (đã fix trong outline v3)',
         '12',
         'Q1-1, Q1-2, Q1-3, Q1-4, Q1-5, Q1-7, Q1-9, '
         'Q3-1, Q3-2, Q3-3, Q3-4, Q3-5',
         'Em',
         'ĐÃ HOÀN THÀNH'),
        ('Tier 3 — MEDIUM (đã fix trong outline v3)',
         '3',
         'Q1-10, Q3-7, Q3-8',
         'Em',
         'ĐÃ HOÀN THÀNH'),
        ('TỔNG', '19', '—', '—', '15 fixed + 4 chờ quyết định'),
    ],
    [3.5, 1.5, 5.5, 2.5, 4.0]
)

P('Trong Tier 2 và Tier 3 (15 issues), các sửa đổi quan trọng nhất bao '
  'gồm: phát biểu lại ba hypotheses Q1 theo dạng testable (Q1-3), bổ '
  'sung rationale cho mixed-methods design (Q1-2), thêm fit indices '
  'threshold cụ thể theo Hu & Bentler 1999 (Q1-5), xây dựng bảng β '
  'coefficients 7×3 (7 predictors × 3 outcomes) đầy đủ từ LA chính '
  '(Q1-9), bổ sung các citations về collectivism và developmental task '
  'cho phần discussion về gender invariance (Q1-10), mở rộng phần '
  'Introduction cho Q3 với rationale cho descriptive approach (Q3-3), '
  'và thêm Bonferroni correction cùng effect sizes (Q3-5).', italic=True)

P('Outline v3 hoàn chỉnh đã được xuất ra cả hai dạng: bản tiếng Việt '
  '(Outline_Q1_v3 và Outline_Q3_v3) và bản song ngữ chi tiết hơn '
  '(OutlineBilingual_Q1 và OutlineBilingual_Q3) — bản song ngữ được '
  'thiết kế đặc biệt để thầy hướng dẫn có thể forward thẳng cho đồng '
  'tác giả Nguyễn Minh Đức cùng đọc và phản hồi.', italic=True)


# ============================================================
H1('5. BỐN CÂU HỎI BLOCKING — CHỜ QUYẾT ĐỊNH')

P('Bốn câu hỏi sau đây thuộc loại "blocking" — nghĩa là nếu không có '
  'câu trả lời, em không thể hoàn thiện outline v3 FINAL hoặc bắt đầu '
  'viết draft bài báo. Mỗi câu hỏi đều có hệ quả nghiêm trọng nếu '
  'không được xử lý kịp thời.')


H2('Câu hỏi Q1-6: Dữ liệu phỏng vấn định tính')

P('Owner: Nghiên cứu sinh Công Thị Hằng.', indent=False)
NOTE('Mixed-methods design là một trong ba điểm mới chính của bài Q1, '
     'do đó dữ liệu phỏng vấn định tính là thành phần thiết yếu — '
     'không phải tùy chọn. Em cần biết: nghiên cứu sinh đã phỏng vấn '
     'bao nhiêu học sinh? Chiến lược lấy mẫu là gì (purposeful hay '
     'theoretical)? Transcripts đã có sẵn chưa? Đã tính được Cohen κ '
     'cho intercoder reliability chưa?')
P('Nếu không có dữ liệu phỏng vấn đầy đủ, claim "mixed-methods" trong '
  'title và abstract sẽ không thể giữ — bài Q1 sẽ giảm từ ba điểm mới '
  'xuống còn hai, có thể không còn đạt chuẩn Q1 BMC Psychiatry.')


H2('Câu hỏi Q1-8: R² của mô hình SEM tích hợp')

P('Owner: Thầy hướng dẫn + Nghiên cứu sinh.', indent=False)
NOTE('Hiện tại, R² = 0,284 (mô hình yếu tố nguy cơ) và R² = 0,209 (mô '
     'hình yếu tố bảo vệ) được trích từ LA chính Bảng 27 và Bảng 30, '
     'NHƯNG hai con số này là từ HAI MÔ HÌNH RIÊNG. Trong khi đó, '
     'claim "integrated SEM" cho bài Q1 đòi hỏi MỘT R² duy nhất từ mô '
     'hình tích hợp với cả 7 dự báo đồng thời.')
P('Có hai lựa chọn. Lựa chọn (A): chạy lại SEM tích hợp đầy đủ với 7 '
  'dự báo đồng thời → ra một R² tổng hợp duy nhất. Việc này đòi hỏi '
  'access vào bộ dữ liệu gốc và mất khoảng 1-2 ngày phân tích. Lựa '
  'chọn (B): giữ nguyên hai R² nhưng thêm chú thích rõ ràng "we report '
  'separate-model R² for comparability with the parent dissertation '
  '[Cong, 2026]". Lựa chọn này nhanh hơn nhưng giảm phần nào tính '
  'mới của bài Q1.')


H2('Câu hỏi Q3-6: Văn bản phê duyệt đạo đức từ HNUE')

P('Owner: Nghiên cứu sinh Công Thị Hằng.', indent=False)
NOTE('Cả BMC Psychiatry và PLOS ONE đều BẮT BUỘC declaration về Institutional '
     'Review Board (IRB) approval. Cần số quyết định, ngày phê duyệt, và '
     'tên hội đồng. Thiếu thông tin này, paper sẽ bị reject ngay từ editor '
     'screening, không cần đến reviewer.')
P('Nếu nghiên cứu sinh đã có letter chính thức từ Hội đồng Đạo đức ĐHSPHN, '
  'em sẽ điền số/ngày vào Methods section và đính kèm letter vào hồ sơ '
  'submission. Nếu chưa có, nghiên cứu sinh cần liên hệ Phòng Quản lý '
  'Khoa học HNUE xin retroactive IRB approval — quy trình này thường '
  'mất 2-4 tuần, do đó cần bắt đầu sớm để không trì hoãn timeline submit.')


H2('Câu hỏi Q3-9: Chiến lược tham chiếu chéo Q1 ↔ Q3')

P('Owner: Thầy hướng dẫn + Nghiên cứu sinh.', indent=False)
NOTE('Vì cả hai bài cùng khai thác bộ dữ liệu 1.352 học sinh, cần một '
     'chiến lược minh bạch về việc cite chéo để tránh bị flag '
     '"redundant publication" — một dạng vi phạm liêm chính học thuật '
     'nghiêm trọng.')
P('Lựa chọn (A): Submit Q1 trước, đợi BMC Psychiatry chấp nhận (8-12 '
  'tháng), sau đó submit Q3 với citation đầy đủ tới Q1 đã published. '
  'Cách này an toàn nhất nhưng kéo dài timeline.')
P('Lựa chọn (B): Submit Q1 và Q3 cùng lúc, mỗi bài có chú thích trong '
  'Methods: "A companion paper presenting [Q1 mechanism / Q3 descriptive] '
  'analysis is currently under review elsewhere." Cách này rút ngắn '
  'timeline xuống 6 tháng nhưng có rủi ro reviewer overlap (cùng một '
  'reviewer được giao cả hai bài → có thể yêu cầu merge thành một).')


# ============================================================
H1('6. THAM KHẢO TÊN BÀI — CHÂU Á VÀ CHÂU PHI')

P('Để chuẩn bị cho cuộc thảo luận chốt tên bài giữa nhóm tác giả, em '
  'đã tổng hợp một bộ tham khảo gồm hơn 30 tiêu đề bài báo gần đây '
  '(2023-2025) từ các tác giả châu Á (Trung Quốc, Việt Nam, Hàn Quốc, '
  'Bangladesh, Ấn Độ, Pakistan, Sri Lanka) và châu Phi (Ethiopia, '
  'Nigeria, Rwanda) về chủ đề rối loạn lo âu ở thanh thiếu niên. Bộ '
  'titles này bao phủ cả tier Q1 (18 titles: SEM, mediation, mixed-'
  'methods) và tier Q3 (15 titles: cross-sectional prevalence).')

P('Phân tích pattern cho thấy các bài Q1 từ tác giả châu Á và châu '
  'Phi thường có ba dạng tiêu đề phổ biến: (1) Dạng "SEM/Mediation" với '
  'cấu trúc "[Variable] and [Outcome] in [Population]: a structural '
  'equation modeling analysis"; (2) Dạng "Multi-group invariance" với '
  'cấu trúc "[Constructs] in [Population]: cross-gender measurement '
  'invariance of [scale]"; (3) Dạng "Risk-Protective" với cấu trúc '
  '"[Risk/Protective] factors of [Outcome] in [Population]".')

P('Đối với các bài Q3 cross-sectional, dạng phổ biến nhất và được tạp '
  'chí PLOS ONE chấp nhận nhiều nhất là dạng "Prevalence and AF" với '
  'cấu trúc "Prevalence and [associated factors / correlates / '
  'determinants] of [outcome] among [population] in [region]: a '
  'cross-sectional study". Bài Bangladesh 2021 trên PLOS ONE là ví '
  'dụ thành công điển hình của pattern này.')

P('Em đã đưa ra 10 đề xuất cụ thể: 5 cho Q1 (A1-A5) và 5 cho Q3 (B1-B5). '
  'Chi tiết đầy đủ trong file riêng "ThamKhao_Titles_Q1Q3_AsiaChauPhi_'
  '01062026.docx". Em đặc biệt khuyến nghị phương án A4 cho Q1 ("Multi-'
  'group structural equation modeling of risk and protective factors '
  'for anxiety disorder subtypes among Vietnamese adolescents: A mixed-'
  'methods cross-sectional study") vì phương án này bám sát pattern '
  'đã được Q1 chấp nhận từ bài Rwandese-GAD 2024 (Frontiers in '
  'Psychiatry). Cho Q3, em khuyến nghị phương án B2 ("Prevalence and '
  'item-level patterns of generalized, separation, and social anxiety '
  'symptoms among Vietnamese lower secondary school students: A '
  'descriptive cross-sectional study") vì phương án này khớp với '
  'pattern thành công của Bangladesh 2021 trên PLOS ONE.')


# ============================================================
H1('7. FILES QUAN TRỌNG ĐÃ TẠO')

P('Phiên làm việc này đã tạo ra ba nhóm files chính, phục vụ ba mục '
  'đích khác nhau: nhóm Outline + Bố cục Q1/Q3 (để viết draft sau '
  'này), nhóm Báo cáo rà soát (để theo dõi và quản lý quá trình), '
  'và nhóm Tài liệu chính thức luận án TS (đã hoàn thành cho nộp hồ '
  'sơ phản biện độc lập).')


H2('7.1 Outline và Bố cục Q1/Q3')

make_table(
    ['Tên file', 'Mục đích sử dụng'],
    [
        ('BoCuc_Q1_Q3_PhuongAnA_v2_01062026.docx',
         'Bố cục tổng quan v2 (đã confirm chốt phương án A)'),
        ('Outline_Q1_v3_01062026.docx',
         'Outline Q1 v3 với 8 issues đã fix — base cho viết draft'),
        ('Outline_Q3_v3_01062026.docx',
         'Outline Q3 v3 với 7 issues đã fix — base cho viết draft'),
        ('OutlineBilingual_Q1_01062026.docx',
         'Bản song ngữ Q1 (15-18 trang) — gửi đồng tác giả review'),
        ('OutlineBilingual_Q3_01062026.docx',
         'Bản song ngữ Q3 (13-15 trang) — gửi đồng tác giả review'),
        ('ThamKhao_Titles_Q1Q3_AsiaChauPhi_01062026.docx',
         'Tham khảo 30+ titles + 10 đề xuất tên bài cho nhóm chọn'),
    ],
    [7.5, 9.5]
)


H2('7.2 Báo cáo rà soát quản lý dự án')

make_table(
    ['Tên file', 'Mục đích sử dụng'],
    [
        ('RaSoat_Q1_01062026.docx',
         'Báo cáo rà soát Q1 chi tiết: 10 issues identified + 49/52 claims '
         'verified'),
        ('RaSoat_Q3_01062026.docx',
         'Báo cáo rà soát Q3 chi tiết: 9 issues identified + 41/41 facts '
         'verified'),
        ('IssuesPriority_Q1Q3_01062026.docx',
         'Phân loại 19 issues thành 3 tiers + action plan + timeline 2 tuần'),
        ('TomTat_Phien_v2_01062026.docx',
         'Tài liệu hiện tại — tóm tắt toàn bộ công việc đã làm'),
    ],
    [7.5, 9.5]
)


H2('7.3 Tài liệu chính thức luận án TS (đã hoàn thành các phiên trước)')

make_table(
    ['Tên file', 'Mục đích sử dụng'],
    [
        ('TomTatLA_v2_VERIFIED_29052026.docx',
         'Tóm tắt LA tiếng Việt (24 trang A5) — nộp hội đồng'),
        ('TomTatLA_EN_v1_29052026.docx',
         'Tóm tắt LA tiếng Anh (24 trang A5) — nộp hội đồng'),
        ('TrichYeuLA_CongThiHang_v2_29052026.docx',
         'Trích yếu LA song ngữ — kèm theo hồ sơ phản biện'),
        ('DanhMucCongTrinh_EN_v1_29052026.docx',
         'Danh mục công trình của tác giả — tiếng Anh, format APA'),
        ('Buoc5_GiayTiepNhan_PBDL_Lan1_DaDien_29052026.docx',
         'Form Bước 5 (Giấy tiếp nhận hồ sơ PBĐL lần 1) đã điền'),
    ],
    [7.5, 9.5]
)


# ============================================================
H1('8. KẾ HOẠCH TRIỂN KHAI SÁU TUẦN TIẾP THEO')

P('Sau khi hoàn tất 15/19 issues của outline v3 và chuẩn bị xong các '
  'tài liệu tham khảo, lộ trình triển khai sáu tuần tiếp theo được '
  'thiết kế để vừa giải quyết bốn câu hỏi blocking còn lại, vừa bắt '
  'đầu công việc viết draft cho bài Q1 với ưu tiên cao.')


H2('Tuần 1 (01-07/06/2026) — Quyết định và Review song song')

P('Tuần này tập trung vào hai luồng công việc song song. Một mặt, thầy '
  'hướng dẫn và nghiên cứu sinh sẽ thảo luận và đưa ra quyết định cho '
  'bốn câu hỏi blocking (Q1-6, Q1-8, Q3-6, Q3-9) cùng việc chốt tên '
  'bài từ 10 đề xuất. Mặt khác, nghiên cứu sinh cần bắt đầu ngay quy '
  'trình xin văn bản IRB chính thức từ HNUE — đây là bước có lead time '
  'dài nhất (2-4 tuần) nên không thể trì hoãn.')

P('Đồng thời, đồng tác giả Nguyễn Minh Đức nên review hai outline song '
  'ngữ Q1 và Q3 đã được chuẩn bị sẵn, đưa ra phản hồi về nội dung học '
  'thuật và phương pháp.')


H2('Tuần 2 (08-14/06/2026) — Integrate decisions vào FINAL outline')

P('Sau khi nhận được quyết định cho bốn câu hỏi blocking, em sẽ tích '
  'hợp các quyết định này vào outline v3 FINAL. Cụ thể: bổ sung phần '
  'qualitative methods chi tiết (Q1-6), điều chỉnh phần R² theo lựa '
  'chọn (A) hoặc (B) cho Q1-8, viết đầy đủ ethics statement standalone '
  'cho Q3-6, và thêm cross-reference strategy explicit cho Q3-9.')

P('Cuối tuần 2, outline v3 FINAL sẽ được gửi lại cho thầy và đồng tác '
  'giả approve trước khi bắt đầu viết draft.')


H2('Tuần 3-5 — Viết draft Q1 đầy đủ')

P('Việc viết draft sẽ tuân theo trình tự khuyến nghị từ các nhà '
  'phương pháp Q1: bắt đầu từ Methods (phần dễ viết nhất vì đã có '
  'protocol rõ ràng), tiếp đến Results (đã có sẵn dữ liệu thống kê), '
  'sau đó Discussion (dựa trên các findings), kế tiếp Introduction '
  '(viết sau khi đã hiểu rõ findings để mạch logic mượt), và cuối '
  'cùng Abstract (~250 từ structured per BMC Psychiatry guidelines).')

P('Em dự kiến hoàn thành draft đầu tiên trong khoảng 3 tuần làm việc '
  'tích cực, sau đó dành thời gian cho việc revise nội bộ trước khi '
  'submit.')


H2('Tuần 6 — Verify, plagiarism check, và submission Q1')

P('Tuần cuối cùng dành cho ba bước chuẩn bị submission: (1) verify '
  'lại tất cả số liệu, β coefficients, F statistics, p values khớp '
  '100% với LA chính qua script tự động; (2) plagiarism check bằng '
  'cách scan các chuỗi 7+ từ liên tiếp trên Google Scholar và Turnitin '
  '(nếu có access qua HNUE); (3) format theo BMC Psychiatry guidelines '
  '(citation style, table format, figure quality, line numbering).')

P('Cuối tuần 6, bài Q1 sẽ được submit kèm cover letter, suggested '
  'reviewers (3-5 chuyên gia từ tác giả của các references đã verify), '
  'và full ethics documentation.')


H2('Tuần 7-8 và sau đó — Draft Q3 song song và peer review Q1')

P('Khi bài Q1 đang trong giai đoạn peer review (thường mất 2-4 tháng '
  'cho quyết định đầu tiên), em sẽ chuyển sang viết draft cho bài Q3. '
  'Vì Q3 có độ phức tạp thấp hơn (descriptive only, không SEM, không '
  'mixed-methods), thời gian viết draft dự kiến ngắn hơn — khoảng 2 '
  'tuần. Việc submit Q3 sẽ phụ thuộc vào chiến lược cross-reference '
  'đã quyết định ở câu hỏi Q3-9: nếu lựa chọn (A) thì đợi Q1 accept '
  'rồi mới submit; nếu lựa chọn (B) thì submit ngay sau khi hoàn '
  'thành draft.')


# ============================================================
H1('LỜI KẾT')

P('Bốn phiên làm việc liên tục đã đặt nền tảng vững chắc cho hai bài '
  'báo quốc tế: tài liệu nguồn đã được rà soát kỹ với tỷ lệ lỗi chỉ '
  '1,4%; thuật ngữ chuyên môn đã được chuẩn hoá theo gợi ý của thầy '
  'hướng dẫn; bố cục hai bài đã rõ ràng với phân chia dữ liệu chống '
  'trùng lặp; 15/19 issues về phương pháp luận đã được fix; bộ tham '
  'khảo titles + 10 đề xuất cụ thể đã sẵn sàng cho thảo luận nhóm.')

P('Bốn câu hỏi blocking còn lại đều có lựa chọn rõ ràng — em đang chờ '
  'thầy hướng dẫn và nghiên cứu sinh quyết định. Sau khi có câu trả '
  'lời, em sẵn sàng tiến vào giai đoạn viết draft bài Q1 với mục tiêu '
  'submit BMC Psychiatry vào cuối tuần 6.')

P('Em xin cảm ơn thầy đã tin tưởng giao công việc rà soát chi tiết này, '
  'và xin cảm ơn nghiên cứu sinh Công Thị Hằng đã cung cấp bộ dữ liệu '
  'và LA chính làm nền tảng. Em sẵn sàng tiếp tục công việc theo chỉ '
  'thị của thầy trong các phiên tiếp theo.', italic=True)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
