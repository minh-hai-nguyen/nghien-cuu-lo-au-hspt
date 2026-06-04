"""Doc trả lời: OR > 1,00 — NAM gấp nữ mấy lần? Format mới (câu trả lời tô xanh trước phụ lục)."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/OR_nam_nu_dien_giai_may_lan.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
BLUE = RGBColor(0, 112, 192)
DARK = RGBColor(31, 73, 125)
GREEN = RGBColor(54, 95, 44)
RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(191, 97, 14)
GRAY = RGBColor(90, 90, 90)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def vn_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p
def blue_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = BLUE
    return p

# ===== TITLE =====
title = d.add_heading('OR > 1,00 trong so sánh nam vs nữ — gấp mấy lần?', level=0)
for r in title.runs: r.font.color.rgb = DARK

sub = d.add_paragraph()
sr = sub.add_run('Diễn giải Odds Ratio đúng cách — phân biệt odds vs nguy cơ thực')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = GRAY
d.add_paragraph()

# ===== CÂU HỎI =====
qbox = d.add_table(rows=1, cols=1); qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]; shade(cell, 'FFF8DC')
p = cell.paragraphs[0]
r = p.add_run('Câu hỏi của thầy ')
r.bold = True
r2 = p.add_run('(em đã sửa chính tả — phần sửa tô xanh)')
r2.italic = True; r2.font.size = Pt(10); r2.font.color.rgb = GRAY
p2 = cell.add_paragraph()
# "i," → "Hỏi:" — tô xanh
r_blue = p2.add_run('Hỏi: ')
r_blue.font.color.rgb = BLUE
r_blue.bold = True
p2.add_run('OR > 1,00 → NAM có khả năng cao hơn nữ mấy lần?')
d.add_paragraph()

# ===== BỐI CẢNH =====
add_h('Bối cảnh', 1)
vn_para('Câu hỏi này thường gặp khi đọc các nghiên cứu dịch tễ tâm lý có hồi quy logistic so sánh giới tính. '
        'Để trả lời chính xác, cần phân biệt 3 khái niệm: ODDS RATIO (OR), RISK RATIO (RR), và "khả năng/nguy cơ" trong tiếng Việt.')
d.add_paragraph()

# ===== KIẾN THỨC NỀN =====
add_h('Kiến thức nền', 1)

vn_para('1. Odds Ratio là gì?', bold=True, size=12)
vn_para('Odds = xác suất xảy ra / xác suất không xảy ra. Ví dụ tỷ lệ lo âu 60 % → odds = 0,60 / 0,40 = 1,5 (cứ 1 người không lo âu thì có 1,5 người lo âu).')
vn_para('OR = odds nhóm 1 / odds nhóm 2. Khi so nam vs nữ: OR = odds_nam / odds_nữ.')
d.add_paragraph()

vn_para('2. Diễn giải OR theo từng giá trị', bold=True, size=12)
or_tbl = d.add_table(rows=8, cols=3); or_tbl.style = 'Table Grid'
hdr = ['OR', 'Diễn giải về ODDS', 'Ý nghĩa thực tế']
for i, h in enumerate(hdr):
    c = or_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
or_data = [
    ('OR = 1,00', 'Odds nam = odds nữ', 'KHÔNG có khác biệt'),
    ('OR = 1,20', 'Odds nam cao hơn nữ 20 %', 'Khác biệt nhỏ'),
    ('OR = 1,50', 'Odds nam cao hơn nữ 50 % (gấp 1,5 lần)', 'Khác biệt vừa'),
    ('OR = 2,00', 'Odds nam GẤP 2 LẦN nữ', 'Khác biệt rõ'),
    ('OR = 3,00', 'Odds nam GẤP 3 LẦN nữ', 'Khác biệt mạnh'),
    ('OR = 5,00', 'Odds nam GẤP 5 LẦN nữ', 'Khác biệt rất mạnh'),
    ('OR < 1,00', 'Odds nam THẤP HƠN nữ (yếu tố BẢO VỆ cho nam)', 'Ngược chiều'),
]
for i, row in enumerate(or_data):
    for j, v in enumerate(row):
        or_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in or_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('Quy tắc: số "lần" = CHÍNH GIÁ TRỊ OR. Cụ thể OR bao nhiêu thì odds gấp bấy nhiêu lần.', bold=True)
d.add_paragraph()

vn_para('3. ⚠ Cảnh báo: "odds" KHÁC "khả năng/nguy cơ"', bold=True, size=12)
warn_tbl = d.add_table(rows=1, cols=1); warn_tbl.style = 'Table Grid'
wc = warn_tbl.rows[0].cells[0]; shade(wc, 'FCE4D6')
wp = wc.paragraphs[0]
wr = wp.add_run('LỖI PHỔ BIẾN: ')
wr.bold = True; wr.font.color.rgb = RED
wp.add_run('Trong tiếng Việt, "khả năng" / "nguy cơ" thường hiểu là PROBABILITY hoặc RISK. '
           'Nhưng OR đo ODDS, KHÔNG ĐO probability hoặc risk trực tiếp. '
           'Khi outcome PHỔ BIẾN (> 20 %), OR PHÓNG ĐẠI nguy cơ thực (RR).')
d.add_paragraph()

vn_para('4. Khi nào OR ≈ RR (đọc "gấp mấy lần" được)?', bold=True, size=12)
rr_tbl = d.add_table(rows=4, cols=2); rr_tbl.style = 'Table Grid'
rr_hdr = ['Tỷ lệ outcome', 'Quan hệ OR vs RR']
for i, h in enumerate(rr_hdr):
    c = rr_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
rr_data = [
    ('Hiếm (< 10 %)', 'OR ≈ RR — có thể đọc "nguy cơ gấp OR lần"'),
    ('Trung bình (10-20 %)', 'OR phóng đại nhẹ — chấp nhận được nhưng cần ghi chú'),
    ('Phổ biến (> 20 %)', 'OR PHÓNG ĐẠI MẠNH — KHÔNG đọc "gấp OR lần" về nguy cơ'),
]
for i, row in enumerate(rr_data):
    for j, v in enumerate(row):
        rr_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in rr_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('5. Ví dụ minh họa cụ thể — OR = 2,00 cho nam vs nữ', bold=True, size=12)
ex_tbl = d.add_table(rows=3, cols=4); ex_tbl.style = 'Table Grid'
ex_hdr = ['Tình huống', 'Tỷ lệ nữ', 'Tỷ lệ nam (tính từ OR=2)', 'RR thực']
for i, h in enumerate(ex_hdr):
    c = ex_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
ex_data = [
    ('Outcome HIẾM', '5 %', '9,6 % (odds 0,053→0,106)', 'RR ≈ 1,9 (gần OR)'),
    ('Outcome PHỔ BIẾN', '40 %', '57 % (odds 0,667→1,333)', 'RR ≈ 1,4 (KHÁC OR)'),
]
for i, row in enumerate(ex_data):
    for j, v in enumerate(row):
        ex_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in ex_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('Trong trường hợp 2: nói "nam nguy cơ gấp 2 lần nữ" sẽ SAI — thực tế chỉ gấp 1,4 lần. '
        'OR = 2 phóng đại RR (1,4) lên thành 2 do outcome 40 % không hiếm.')
d.add_paragraph()

vn_para('6. Đối chiếu với các bài trong CSDL của em', bold=True, size=12)
db_tbl = d.add_table(rows=4, cols=3); db_tbl.style = 'Table Grid'
db_hdr = ['Bài', 'OR/AOR', 'Diễn giải']
for i, h in enumerate(db_hdr):
    c = db_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
db_data = [
    ('QT005 Zhu 2025 TQ', 'AOR = 13,71 (ngủ < 5h)', 'Odds gấp 13,7 lần. Outcome ~14 % nên AOR phóng đại nhẹ RR (~10-12).'),
    ('QT031 Islam 2025 (59 nước)', 'AOR = 1,51 (nữ vs nam)', 'Odds nữ gấp 1,5 lần nam. Outcome 20-30 % nên RR ≈ 1,3-1,4 (phóng đại nhẹ).'),
    ('QT017 Puyat 2025 Philippines', 'aPR = 1,60', 'Đây là PREVALENCE RATIO (không phải OR) — đọc TRỰC TIẾP "nguy cơ gấp 1,6 lần".'),
]
for i, row in enumerate(db_data):
    for j, v in enumerate(row):
        db_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in db_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

# =====================================================
# CÂU TRẢ LỜI (TÔ XANH, GOM LẠI TRƯỚC PHỤ LỤC)
# =====================================================
add_h('Câu trả lời', 1, BLUE)

blue_para('Đáp án ngắn — số "lần" chính là giá trị OR:', bold=True, size=12)
d.add_paragraph()

ans_tbl = d.add_table(rows=4, cols=2); ans_tbl.style = 'Table Grid'
ans_hdr = ['OR', 'Nam gấp nữ bao nhiêu LẦN (về odds)']
for i, h in enumerate(ans_hdr):
    c = ans_tbl.rows[0].cells[i]; shade(c, 'BDD7EE')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = BLUE
ans_data = [
    ('OR = 1,5', '1,5 lần (cao hơn 50 %)'),
    ('OR = 2,0', '2 lần'),
    ('OR = 3,0', '3 lần'),
]
for i, row in enumerate(ans_data):
    for j, v in enumerate(row):
        c = ans_tbl.rows[i+1].cells[j]
        pp = c.paragraphs[0]
        if j == 1: pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        rr = pp.add_run(v); rr.font.color.rgb = BLUE; rr.bold = (j == 0)
d.add_paragraph()

blue_para('Diễn giải CHÍNH XÁC (luôn đúng):', bold=True, size=12)
blue_para('"Odds [outcome] của nam cao gấp OR lần so với nữ"', bold=True)
blue_para('Đây là cách an toàn nhất, không bị phóng đại.')
d.add_paragraph()

blue_para('Cảnh báo: KHÔNG nói "nam có nguy cơ gấp OR lần" trừ khi:', bold=True, size=12)
for it in [
    'Outcome HIẾM (< 10 %) — khi đó OR ≈ RR, đọc "nguy cơ gấp OR lần" chấp nhận được',
    'Outcome PHỔ BIẾN (> 20 %, như lo âu HS VN 40-60 %) — OR phóng đại nguy cơ, KHÔNG đọc "gấp OR lần" về nguy cơ',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Quy tắc 3 bước cho thầy đọc paper:', bold=True, size=12)
for it in [
    '1. Tìm tỷ lệ outcome trong nhóm tham chiếu (nữ). Vd: paper ghi "lo âu nữ 25 %"',
    '2. Nếu < 10 % → OR ≈ RR, đọc trực tiếp "gấp OR lần"',
    '3. Nếu ≥ 20 % → tránh nói "nguy cơ gấp OR lần". Nói "ODDS gấp OR lần" hoặc tự convert sang RR',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Ghi nhớ 1 dòng:', bold=True, size=12)
blue_para('OR > 1,00 → ODDS nam gấp OR lần nữ. Số "lần" = chính giá trị OR. '
          'Chỉ nói "NGUY CƠ gấp OR lần" khi outcome HIẾM (< 10 %) — outcome phổ biến thì OR phóng đại.', bold=True)

d.add_paragraph()

# =====================================================
# PHỤ LỤC
# =====================================================
add_h('Phụ lục — Tài liệu tham khảo', 1)
refs = [
    'Szumilas M. (2010). Explaining odds ratios. Journal of the Canadian Academy of Child and Adolescent Psychiatry, 19(3):227-229.',
    'Zhang J, Yu KF. (1998). What\'s the relative risk? A method of correcting the odds ratio in cohort studies of common outcomes. JAMA, 280(19):1690-1691.',
    'Hosmer DW, Lemeshow S, Sturdivant RX. (2013). Applied Logistic Regression (3rd ed.). Wiley.',
    'Zhu et al. (2025). [QT005 — AOR=13,71 cho ngủ<5h]. BMC Public Health.',
    'Islam et al. (2025). [QT031 — AOR=1,51 cho nữ vs nam, 59 LMICs]. Journal of Affective Disorders.',
    'Puyat et al. (2025). [QT017 — aPR=1,60, sử dụng PREVALENCE RATIO không phải OR]. Phil J Health.',
    'Doc liên quan trong dự án: 01_Bao-cao/AOR_Adjusted_Odds_Ratio_cach_tinh.docx (chi tiết về AOR + công thức β → AOR).',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs: r.font.size = Pt(10)

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
