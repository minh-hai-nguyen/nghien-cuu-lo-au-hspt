# -*- coding: utf-8 -*-
"""Dịch 4 bài paywall — chỉ có abstract từ Crossref/Semantic Scholar (chưa có PDF)"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                t += ' ' + cell.text
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

def disclaimer(doc):
    p = doc.add_paragraph()
    r = p.add_run('⚠ LƯU Ý: Bản dịch này CHỈ DỰA TRÊN ABSTRACT từ Crossref/Semantic Scholar — chưa có PDF đầy đủ (paywall). Cần truy cập thư viện ĐH để có toàn văn.')
    r.font.name='Times New Roman'; r.font.size=Pt(11); r.bold=True; r.italic=True
    r.font.color.rgb = RGBColor(0xCC, 0x66, 0x00)

# ========== A4 Chen 2025 — JAffectDisord COVID Meta ==========
print('A4 Chen 2025 (abstract-only)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.01.029', size=10)
add_heading(doc, 'Yếu tố bảo vệ và yếu tố nguy cơ của lo âu ở trẻ em và vị thành niên trong COVID-19: Tổng quan hệ thống và phân tích tổng hợp ba cấp', 1)
h = doc.add_paragraph()
r = h.add_run('Protective and risk factors of anxiety in children and adolescents during COVID-19: A systematic review and three level meta-analysis')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True
disclaimer(doc)

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'Huijing Chen, Qi Wang, Jiangle Zhu, Yi Zhu, Feixu Yang, Junyi Hui, Xinfeng Tang, Tianming Zhang'),
    ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,5) — Elsevier'),
    ('Xuất bản', '2025, Vol. 374, pp. 408–432, 25 trang'),
    ('DOI', '10.1016/j.jad.2025.01.029'),
    ('Loại NC', 'Tổng quan hệ thống + Phân tích tổng hợp 3 cấp (three-level meta-analysis)'),
    ('Mẫu', '141 nghiên cứu, hơn 1.000.000 trẻ em + VTN'),
    ('Phạm vi', 'Lo âu trẻ em + VTN trong đại dịch COVID-19 toàn cầu'),
    ('Truy cập', '🔒 Paywall — ScienceDirect'),
])
add_page_ref(doc, '408–432', 'J Affect Disord', 'Vol. 374, 2025')

add_heading(doc, 'TÓM LƯỢC NỘI DUNG (từ abstract)', 2)
add_p(doc, 'Bài tổng quan + meta-analysis 3 cấp này phân tích các yếu tố ảnh hưởng đến lo âu ở trẻ em và VTN trong đại dịch COVID-19. Phân tích 141 NC bao phủ HƠN 1 TRIỆU người tham gia. Đã xác định 14 yếu tố BẢO VỆ và 29 yếu tố NGUY CƠ trong nhiều domain khác nhau.')

p = doc.add_paragraph()
r = p.add_run('Phát hiện chính: "Chức năng cảm xúc" (emotional functioning) là yếu tố BẢO VỆ MẠNH. Trong khi "nghiện thiết bị điện tử/internet" và "mức độ lo âu của người chăm sóc" là yếu tố NGUY CƠ. Hỗ trợ gia đình và nguồn lực cộng đồng là các bộ đệm (buffer) quan trọng.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_p(doc, 'Kết luận: Cần ưu tiên các chiến lược phòng ngừa và can thiệp cho trẻ em + VTN nguy cơ cao, đặc biệt qua chương trình cộng đồng trong các cuộc khủng hoảng y tế công cộng.')

add_red_heading(doc, 'GIÁ TRỊ CHO ĐỀ TÀI VN')
add_red(doc, '• META 141 NC + 1M+ trẻ — bằng chứng MẠNH NHẤT về yếu tố nguy cơ/bảo vệ COVID era. Q1 IF~6,5.')
add_red(doc, '• Yếu tố bảo vệ KÉP: emotional functioning + family support + community resources — phù hợp đề cương VN.')
add_red(doc, '• "Caregiver anxiety" là yếu tố nguy cơ → gợi ý can thiệp song song cha mẹ.')
add_red(doc, '• Internet/device addiction → phù hợp Zheng A6 (MXH+lo âu) và Nature SM 2025 (QT27).')
add_red(doc, '• 29 nguy cơ vs 14 bảo vệ → mất cân bằng — VN cần tăng cường yếu tố bảo vệ.')
add_red(doc, '• Hạn chế: chưa có toàn văn → không biết chi tiết effect sizes, OR/RR cụ thể từng yếu tố.')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐⭐ Rất cao (chỉ tính abstract). Cần tải full PDF từ thư viện ĐH để rút trích chi tiết.', bold=True)
save(doc, '55_Chen_COVID_141Studies_2025_ABSTRACT.docx')

# ========== B1 Zhang 2025/2026 — Long-Term CBT Bayesian MA ==========
print('B1 Zhang Bayesian MA (abstract-only)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1002/jclp.70069', size=10)
add_heading(doc, 'Hiệu quả lâu dài của CBT tại trường ở trẻ em và vị thành niên nguy cơ thấp: Phân tích tổng hợp Bayesian', 1)
h = doc.add_paragraph()
r = h.add_run('Long-Term Effects of School-Based CBT in Low-Risk Children and Adolescents: A Bayesian Meta-Analysis')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True
disclaimer(doc)

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'Xiangying Zhang, Zhide Liang, Jihye Kang'),
    ('Tạp chí', 'Journal of Clinical Psychology (Q1) — Wiley'),
    ('Xuất bản', 'Online 11/2025, in print Vol. 82, pp. 248–259, 03/2026'),
    ('DOI', '10.1002/jclp.70069'),
    ('Loại NC', 'Phân tích tổng hợp Bayesian — 31 RCT'),
    ('Mẫu', '31 RCT, n = 19.865 trẻ + VTN nguy cơ thấp'),
    ('Phạm vi', 'CBT phổ quát tại trường — trầm cảm + lo âu, follow-up dài hạn'),
    ('Truy cập', '🔒 Paywall — Wiley'),
])
add_page_ref(doc, '248–259', 'J Clin Psychol', 'Vol. 82(3), 2026')

add_heading(doc, 'TÓM LƯỢC NỘI DUNG (từ abstract)', 2)
add_p(doc, 'Phân tích tổng hợp Bayesian 31 RCT về CBT phổ quát tại trường cho trầm cảm và lo âu ở thanh thiếu niên nguy cơ thấp (n = 19.865). Tổng số mẫu lớn nhất hiện nay cho can thiệp CBT trường.')

p = doc.add_paragraph()
r = p.add_run('Kết quả: Cải thiện CÓ Ý NGHĨA THỐNG KÊ nhưng KHIÊM TỐN trên triệu chứng trầm cảm; giảm NHỎ trên triệu chứng lo âu. Hiệu quả DUY TRÌ tới 1 năm sau can thiệp.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

p = doc.add_paragraph()
r = p.add_run('Kết luận tác giả: "Chất lượng bằng chứng nền rất thấp khiến phát hiện này CHƯA ĐỦ vững chắc để hỗ trợ triển khai rộng rãi tại thời điểm này." Tức là: CBT trường có hứa hẹn nhưng cần RCT chất lượng cao hơn trước khi mở rộng.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_red_heading(doc, 'GIÁ TRỊ CHO ĐỀ TÀI VN')
add_red(doc, '• Bayesian MA 31 RCT, n=19.865 — tổng mẫu lớn nhất. Q1 Wiley.')
add_red(doc, '• Phát hiện QUAN TRỌNG: CBT phổ quát có hiệu quả NHỎ và CHẤT LƯỢNG NỀN THẤP — TRÁI với QT29 BMC NMA 2025 (CBT SUCRA 0,66 hạng 2).')
add_red(doc, '• Mâu thuẫn này phản ánh khác biệt giữa CBT PHỔ QUÁT (low-risk, dilution) vs CBT CÓ MỤC TIÊU (high-risk).')
add_red(doc, '• Phù hợp B5 UK (mindfulness 8.376 HS thất bại do engagement) — universal interventions thường yếu hiệu quả.')
add_red(doc, '• Gợi ý cho VN: nên thử CBT TARGETED (HS có triệu chứng) thay vì universal — giảm dilution.')
add_red(doc, '• Hiệu quả duy trì 1 năm — đáng khích lệ, nhưng cần RCT VN dài hạn.')
add_red(doc, '• Hạn chế: chưa có full PDF → chưa biết SMD cụ thể, heterogeneity I².')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. Bayesian MA Q1 — quan trọng cho cảnh báo về universal CBT.', bold=True)
save(doc, '56_Zhang_Bayesian_CBT_2026_ABSTRACT.docx')

# ========== B4 Qiaochu 2025 — Mobile CBT SR ==========
print('B4 Qiaochu Mobile CBT (abstract-only)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1002/cpp.70173', size=10)
add_heading(doc, 'Hiệu quả của CBT trên thiết bị di động cho trầm cảm và lo âu ở trẻ em và người trẻ: Tổng quan hệ thống các thử nghiệm ngẫu nhiên có đối chứng', 1)
h = doc.add_paragraph()
r = h.add_run('Effectiveness of Mobile-Based Cognitive Behavioural Therapy for Depression and Anxiety in Children and Young People: A Systematic Review of Randomized Controlled Trials')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True
disclaimer(doc)

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'Zhang Qiaochu, Wang Yahui'),
    ('Tạp chí', 'Clinical Psychology & Psychotherapy (Q1) — Wiley'),
    ('Xuất bản', '2025, Vol. 32(6), e70173'),
    ('DOI', '10.1002/cpp.70173'),
    ('Loại NC', 'Tổng quan hệ thống 9 RCT'),
    ('Mẫu', '9 RCT, N = 2.479 trẻ em + người trẻ'),
    ('Phạm vi', 'CBT qua MOBILE APP cho trầm cảm + lo âu trẻ em/VTN'),
    ('Truy cập', '🔒 Paywall — Wiley (OA bronze)'),
])
add_page_ref(doc, 'e70173', 'Clin Psychol Psychother', 'Vol. 32(6), 2025')

add_heading(doc, 'TÓM LƯỢC NỘI DUNG (từ abstract)', 2)
add_p(doc, 'Tổng quan hệ thống 9 RCT đánh giá hiệu quả của CBT cung cấp qua MOBILE APP cho trầm cảm và lo âu ở trẻ em + người trẻ. Tổng cộng N = 2.479 người tham gia.')

p = doc.add_paragraph()
r = p.add_run('Kết quả chính: 7/8 NC (87,5%) đo trầm cảm cho thấy GIẢM CÓ Ý NGHĨA triệu chứng trầm cảm. NHƯNG bằng chứng cho lo âu HẠN CHẾ — chỉ 2/6 NC cho thấy hiệu quả có ý nghĩa.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_p(doc, 'Kết luận: Mobile CBT có HỨA HẸN cho TRẦM CẢM ở trẻ em + người trẻ. Các can thiệp dễ tiếp cận này có thể giúp lấp khoảng trống dịch vụ SKTT, ĐẶC BIỆT trong bối cảnh hạn chế nguồn lực.')

add_red_heading(doc, 'GIÁ TRỊ CHO ĐỀ TÀI VN')
add_red(doc, '• 9 RCT, N=2.479 — đủ lớn để rút kết luận về mobile CBT cho VTN. Q1 Wiley.')
add_red(doc, '• Phát hiện QUAN TRỌNG: hiệu quả TRẦM CẢM (7/8) nhưng LO ÂU YẾU (2/6) — Mobile CBT KHÔNG MẠNH cho lo âu.')
add_red(doc, '• Mâu thuẫn với B11 Nhật iCBT subthreshold SAD (RCT đa trung tâm dương tính cho SAD).')
add_red(doc, '• Phù hợp B2 JMIR Digital SAD MA 2025 (21 RCT — digital có hiệu quả TB cho SAD).')
add_red(doc, '• Gợi ý: VN nên thử mobile CBT cho TRẦM CẢM trước (bằng chứng mạnh) hoặc kết hợp.')
add_red(doc, '• "Resource-limited settings" → trực tiếp phù hợp VN.')
add_red(doc, '• Hạn chế: chưa có PDF → chưa biết các app cụ thể, dropout rates, age range chi tiết.')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. SR Q1, phù hợp đề tài LMIC, kết quả khác biệt depression vs anxiety quan trọng.', bold=True)
save(doc, '57_Qiaochu_MobileCBT_SR_2025_ABSTRACT.docx')

# ========== B10 Menon 2025 — APJPH Scoping Review LMIC East Asia Pacific ==========
print('B10 Menon LMIC SEA Pacific (abstract-only)...')
doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1177/10105395241313154', size=10)
add_heading(doc, 'Các can thiệp được đánh giá nhắm vào sức khỏe tâm thần và sức khỏe tâm lý xã hội của trẻ em và vị thành niên: Tổng quan phạm vi tập trung vào các nước thu nhập thấp và trung bình ở Đông Á và Thái Bình Dương', 1)
h = doc.add_paragraph()
r = h.add_run('Evaluated Interventions Targeting the Mental Health and Psychosocial Wellbeing of Children and Adolescents: A Scoping Review Focused on Low- and Middle-Income Countries in East Asia and the Pacific')
r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True
disclaimer(doc)

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tác giả', 'Vinay Menon, Miika Coppard, Samuel McEwen, Lorena Romero, Elissa Kennedy, Peter Azzopardi'),
    ('Tạp chí', 'Asia Pacific Journal of Public Health (Q2) — SAGE'),
    ('Xuất bản', '2025, Vol. 37(4), pp. 332–346, 15 trang'),
    ('DOI', '10.1177/10105395241313154'),
    ('Loại NC', 'Tổng quan phạm vi (Scoping Review)'),
    ('Mẫu', '69 nghiên cứu / 12 quốc gia LMIC Đông Á + TBD: 32 RCT, 31 NC trước-sau, 6 đánh giá sau can thiệp'),
    ('Phạm vi', 'Can thiệp SKTT + tâm lý xã hội cho trẻ em + VTN ở LMIC Đông Á + Thái Bình Dương'),
    ('Truy cập', '🔒 Paywall — SAGE'),
])
add_page_ref(doc, '332–346', 'Asia Pac J Public Health', 'Vol. 37(4), 2025')

add_heading(doc, 'TÓM LƯỢC NỘI DUNG (từ abstract)', 2)
add_p(doc, 'Tổng quan phạm vi xem xét các can thiệp SKTT + tâm lý xã hội cho trẻ em + VTN tại LMIC Đông Á + Thái Bình Dương. Tổng cộng 69 NC từ 12 quốc gia: 32 RCT, 31 NC trước-sau, 6 đánh giá sau can thiệp.')

p = doc.add_paragraph()
r = p.add_run('Phát hiện chính: Có sự TẬP TRUNG MẤT CÂN ĐỐI vào "năng lực cá nhân" (individual capacity) — tức can thiệp PHÒNG NGỪA và "quản lý lâm sàng" (clinical management). KHOẢNG TRỐNG ở: thúc đẩy SK dựa vào CỘNG ĐỒNG, phòng ngừa ở CẤP GIA ĐÌNH, và dịch vụ ĐÁP ỨNG dài hạn.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

p = doc.add_paragraph()
r = p.add_run('Phân bố địa lý: phần lớn NC tập trung ở TRUNG QUỐC + Đông Nam Á. Các nước nhỏ hơn và vùng Thái Bình Dương có ĐẠI DIỆN TỐI THIỂU.')
r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

add_red_heading(doc, 'GIÁ TRỊ CHO ĐỀ TÀI VN')
add_red(doc, '• 69 NC / 12 nước LMIC Đông Á + TBD — bao gồm VN. Phản ánh KHU VỰC trực tiếp cho đề tài.')
add_red(doc, '• KHOẢNG TRỐNG xác định rõ: thiếu can thiệp CỘNG ĐỒNG + GIA ĐÌNH + DỊCH VỤ DÀI HẠN.')
add_red(doc, '• Phù hợp với phát hiện Dong A5 (TQ): kênh giao tiếp gia đình OR=0,22 — gia đình là yếu tố bảo vệ chính.')
add_red(doc, '• Phù hợp với UNICEF VN22 + đề cương VN giai đoạn 2: cần can thiệp PHỐI HỢP gia đình-trường-cộng đồng.')
add_red(doc, '• Cảnh báo: TQ + ĐNA chiếm phần lớn → có thể thiếu chú trọng VN cụ thể.')
add_red(doc, '• Gợi ý cho đề cương VN: ƯU TIÊN can thiệp CỘNG ĐỒNG + GIA ĐÌNH thay vì chỉ trường — sẽ lấp khoảng trống.')
add_red(doc, '• Hạn chế: chưa có PDF → chưa biết VN có bao nhiêu NC, loại nào.')
add_p(doc, 'Đánh giá: ⭐⭐⭐⭐ Cao. Scoping review khu vực ĐN Á LMIC — trực tiếp xác định khoảng trống cho VN.', bold=True)
save(doc, '58_Menon_LMIC_SEA_Pacific_2025_ABSTRACT.docx')

print('DONE 4 paywall abstracts')
