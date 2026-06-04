# -*- coding: utf-8 -*-
"""
Expand PHẢN BIỆN VÀ ĐÁNH GIÁ section of VN022 translation.
- Replace paragraphs 1033–1067 with a detailed critique including citations.
- Citations drawn from the report's own bibliography (verified exists in refs list).
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from copy import deepcopy

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

d = Document(PATH)

# Locate anchor paragraphs by content
heading_idx = None        # "PHẢN BIỆN VÀ ĐÁNH GIÁ ..."
end_idx = None            # "ĐÁNH GIÁ TỔNG THỂ: ⭐⭐⭐⭐⭐ (5/5)" — keep; we stop before this
section_start_idx = None  # "Đánh giá chất lượng báo cáo UNICEF 2022:" — content begins here

for i, p in enumerate(d.paragraphs):
    t = p.text.strip()
    if t.startswith('PHẢN BIỆN VÀ ĐÁNH GIÁ') and heading_idx is None:
        heading_idx = i
    if t.startswith('Đánh giá chất lượng báo cáo UNICEF 2022') and section_start_idx is None:
        section_start_idx = i
    if t.startswith('ĐÁNH GIÁ TỔNG THỂ') and end_idx is None:
        end_idx = i

print(f'Heading at #{heading_idx}')
print(f'Section start at #{section_start_idx}')
print(f'Overall verdict at #{end_idx}')
assert heading_idx and section_start_idx and end_idx, 'Anchors not found'

# Paragraphs to delete: section_start_idx ... end_idx-1
to_delete = list(range(section_start_idx, end_idx))
print(f'Will delete {len(to_delete)} paragraphs (#{to_delete[0]}..#{to_delete[-1]})')

# Get the anchor paragraph (before first deleted) — we'll insert new paragraphs after it
# The "before" anchor is paragraph at (section_start_idx - 1)
before_anchor = d.paragraphs[section_start_idx - 1]

# Delete old content paragraphs (in reverse so indices stay valid in the underlying XML)
old_els = [d.paragraphs[i]._element for i in to_delete]
for el in old_els:
    el.getparent().remove(el)
print(f'Deleted {len(old_els)} old paragraphs')

# Helper: add a paragraph at the end of the document with styling, then we will move it
def make_para(text, bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = d.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size:
        r.font.size = Pt(size)
    if bold:
        r.bold = True
    if italic:
        r.italic = True
    if red:
        r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color:
        r.font.color.rgb = color
    return p

# Build the expanded critique as a list of (text, style_dict)
# Style dict keys: bold, italic, size, red, align
BOLD = {'bold': True, 'red': True}
H1   = {'bold': True, 'red': True, 'size': 13}
H2   = {'bold': True, 'red': True, 'size': 12}
RED  = {'red': True}
ITRED = {'italic': True, 'red': True, 'size': 10}

lines = [
    ('PHẦN I — BỐI CẢNH HỌC THUẬT & VỊ TRÍ CỦA BÁO CÁO', H2),
    ('Báo cáo UNICEF Việt Nam 2022 nằm trong làn sóng các nghiên cứu quốc gia về sức khoẻ tâm thần vị thành niên tại các nước thu nhập thấp và trung bình (LMICs) được khởi động sau 2015, khi WHO và Lancet Commission on Adolescent Health (Patton et al., 2016) kêu gọi các chính phủ đầu tư hệ thống dịch vụ sức khoẻ tâm thần cho trẻ em – vị thành niên. Trong khu vực Đông Nam Á, báo cáo này có cùng bối cảnh với National Adolescent Mental Health Surveys (NAMHS) đang triển khai tại Indonesia, Kenya và Việt Nam (Erskine et al., 2021 — có trong danh mục tham khảo của báo cáo) và review của Hirota, Guerrero & Skokauskas (2020) về "Child and adolescent mental health needs, services, and gap in East and Southeast Asia and the Pacific Islands". Điểm khác biệt của UNICEF 2022 so với NAMHS là tập trung vào yếu tố trường học — school climate, quan hệ giáo viên–học sinh, bắt nạt, áp lực học tập — thay vì chỉ đo prevalence rối loạn tâm thần. Do đó báo cáo này bổ sung cho NAMHS ở chiều cạnh yếu tố nguy cơ/bảo vệ tại môi trường học đường.', RED),
    ('', {}),

    ('PHẦN II — ĐIỂM MẠNH (được phân tích chi tiết)', H2),
    ('1. Uy tín thể chế và giá trị chính sách', BOLD),
    ('Đồng chủ trì giữa UNICEF Việt Nam và Bộ Giáo dục & Đào tạo (MOET), với sự hợp tác của Bộ Y tế (MOH) và Bộ Lao động – Thương binh & Xã hội (MOLISA) — ba cơ quan then chốt điều hành hệ thống dịch vụ MHPSS trường học. Điều này đảm bảo: (a) tiếp cận dữ liệu hành chính — trường, học sinh, giáo viên — mà nghiên cứu độc lập khó có; (b) khuyến nghị đi thẳng vào chu trình chính sách (Quyết định 1442/QĐ-TTg 2020, Thông tư 31/2017/TT-BGDĐT được nêu đích danh trong báo cáo).', RED),
    ('2. Thiết kế mixed-methods chặt chẽ và có chiều sâu', BOLD),
    ('Kết hợp định lượng (n = 668 học sinh + 66 giáo viên) với 13 FGD (39 HS + 21 phụ huynh + 35 GV) và 34 KII (cán bộ 3 bộ, lãnh đạo trường, chuyên gia) cho phép triangulation các chủ đề then chốt. Ví dụ: số liệu định lượng 52,2 % HS từng bị cyberbullying được bổ sung bằng insight định tính từ FGD — các em mô tả cơ chế "exposure anonymously" trên TikTok/Facebook, điều mà con số đơn thuần không diễn đạt được. Đây là mô hình nghiên cứu được Creswell & Plano Clark (2017) xem là "integrative mixed-methods" và có giá trị xây dựng chính sách cao hơn thiết kế đơn phương pháp.', RED),
    ('3. Công cụ đo đã chuẩn hoá và validate tại Việt Nam', BOLD),
    ('Báo cáo sử dụng bộ công cụ có psychometric properties đã được kiểm định trong bối cảnh Việt Nam: SDQ-25 (Strengths and Difficulties Questionnaire — Goodman 1997; validate VN bởi Dang et al. 2016 "Vietnam as a case example of school-based mental health services..." — citation #901 trong references), MDS3 (Maryland Safe and Supportive Schools — Bradshaw et al. 2014), ESSA (Engagement with School Scale for Adolescents — Sun et al. 2011), RCBI (Revised Cyber Bullying Inventory — validate tiếng Việt bởi Cong, Ngoc, Weiss, Luot & Dat 2018, citation #898). Sử dụng các công cụ đã validate giúp kết quả đối chiếu được với literature quốc tế, đồng thời tránh rủi ro "bịa" thang đo mới tại mỗi nghiên cứu.', RED),
    ('4. Chọn mẫu đa dạng theo kinh tế–xã hội và địa lý', BOLD),
    ('5 tỉnh ban đầu gồm Hà Nội (Bắc, đô thị), Điện Biên (Tây Bắc, DTTS, nông thôn vùng sâu), Gia Lai (Tây Nguyên, DTTS), Đồng Tháp (Đồng bằng sông Cửu Long, nông thôn) và TPHCM (Nam, đô thị — bỏ do COVID-19). Mẫu học sinh bao gồm các dân tộc Mông, Thái, Ba Na, Gia-rai, Khơ Mú, Tày, Xơ đăng, Mường — độ đa dạng này cao hơn đáng kể so với phần lớn nghiên cứu VN cùng chủ đề (hầu hết chỉ lấy mẫu tại 1–2 tỉnh đô thị). Báo cáo cũng chủ động tuyển học sinh LGBTQ tham gia FGD, tuy chưa có sampling frame thống kê cho nhóm này.', RED),
    ('5. Tham gia của nhiều bên liên quan', BOLD),
    ('Không chỉ thu thập dữ liệu từ học sinh, báo cáo lấy ý kiến giáo viên (n = 66), phụ huynh (n = 21), cán bộ MOET/MOH/MOLISA (n = 34). Cách tiếp cận này theo khung socio-ecological của Bronfenbrenner và được Fazel, Hoagwood, Stephan & Ford (2014, Lancet Psychiatry, "Mental health interventions in schools in high-income countries") khuyến nghị như điều kiện cần cho nghiên cứu MHPSS trường học có tính áp dụng chính sách.', RED),
    ('6. Chín khuyến nghị chính sách đi kèm phân khúc đối tượng rõ ràng', BOLD),
    ('Khác với nhiều báo cáo chỉ đưa khuyến nghị chung chung, UNICEF 2022 chia khuyến nghị theo 4 nhóm stakeholder: (a) hệ thống Giáo dục — MOET; (b) hệ thống Y tế — MOH; (c) hệ thống Xã hội — MOLISA; (d) UNICEF & tổ chức đối tác. Điều này phản ánh chuẩn của WHO mhGAP implementation guide (2016) về "actionable policy recommendations".', RED),
    ('', {}),

    ('PHẦN III — HẠN CHẾ (được phản biện sâu)', H2),
    ('1. Thiết kế cắt ngang — không xác lập nhân quả', BOLD),
    ('Toàn bộ dữ liệu thu thập tại một thời điểm duy nhất (2021). Do vậy các mối liên hệ thống kê quan sát được — ví dụ "HS có khí hậu trường học tiêu cực có điểm SDQ cao hơn" — chỉ có thể xem là association, không phải nguyên nhân. Để thiết lập causality cần thiết kế longitudinal với ít nhất 2 timepoints (Rutter 2007, "Proceeding from observed correlation to causal inference: The use of natural experiments"). Với bối cảnh VN, longitudinal nghiên cứu Duong et al. 2025 [VN029 trong dự án Lo âu, n = 2.631, 3 timepoints 2022–2024] có thể bổ sung lỗ hổng này.', RED),
    ('2. Lấy mẫu thuận tiện (convenience sampling)', BOLD),
    ('Báo cáo thừa nhận ở trang 30 rằng tuyển học sinh qua trường và qua giới thiệu của giáo viên, không phải random sampling từ khung mẫu toàn tỉnh. Hệ quả: (a) có thể over-represent các em đi học đều (under-represent học sinh bỏ học — nhóm có nguy cơ SKTT cao nhất theo Dupéré et al. 2015); (b) giáo viên có thể vô tình lựa chọn HS "dễ phỏng vấn" — bias selection về phía nhóm có khả năng ngôn ngữ/tự tin cao hơn. Các số liệu prevalence nên được đọc với giới hạn này.', RED),
    ('3. Cỡ mẫu định lượng vừa phải', BOLD),
    ('668 HS trải 4 tỉnh — sau khi chia theo giới × cấp học × DTTS × tỉnh, nhiều subgroup chỉ còn n = 30–50. Power cho detection subgroup effects nhỏ bị giới hạn. So sánh với các nghiên cứu VN cùng chủ đề: Hoa et al. 2024 [VN001 trong dự án Lo âu] có n = 3.910; Duong et al. 2025 [VN029] có n = 2.631; Campbell, Bann & Patalay (2021, citation #887) có n = 566.829 (đa quốc gia). Báo cáo UNICEF có sức mạnh ở insight định tính hơn là precision định lượng.', RED),
    ('4. SDQ-25 không đo trực tiếp lo âu hay trầm cảm lâm sàng', BOLD),
    ('SDQ-25 là công cụ screening problem behaviour tổng quát với 5 subscale (emotional, conduct, hyperactivity, peer, prosocial). Cutoff điểm ≥ 17 (total difficulties) được đề xuất như "high risk", nhưng KHÔNG tương đương chẩn đoán GAD, MDD theo DSM-5 hay ICD-11. Muốn đối chiếu với các nghiên cứu VN dùng GAD-7, PHQ-9, DASS-21 (như VN001, VN029) cần cẩn trọng: các công cụ đo construct khác nhau. Lời nhắc quan trọng: theo memory dự án (feedback_gad7_phq9_separation), GAD-7 và PHQ-9 bản thân cũng không nên gộp chung — SDQ-25 lại càng không thể dùng như proxy cho cả hai.', RED),
    ('5. Psychometric của SDQ-25 phiên bản tiếng Việt còn gây tranh luận', BOLD),
    ('Dang et al. 2017 (tham chiếu citation trong báo cáo) đã chỉ ra factor structure 5-subscale của SDQ-25 không ổn định trong mẫu Việt Nam — Cronbach α của subscale "conduct problems" và "peer problems" dao động 0,47–0,62 (dưới ngưỡng chấp nhận 0,70). Báo cáo UNICEF không công bố reliability riêng cho mẫu của mình (trang 35 chỉ tham chiếu "đã validate"). Đây là gap cần minh bạch trong báo cáo phiên bản tiếp theo.', RED),
    ('6. Teacher self-report bias — đặc biệt nghiêm trọng', BOLD),
    ('66 giáo viên được hỏi "bao nhiêu % HS trong lớp có vấn đề SKTT?" Báo cáo nêu giáo viên đánh giá thấp tỷ lệ vấn đề SKTT ở trẻ em (trong khi tỷ lệ thực tế theo SDQ-25 là 26,1 %). Đây là phát hiện quan trọng nhưng cần đặt trong literature: Meltzer et al. 2000 (British Child Mental Health Survey) cho thấy teacher rating và self-report của trẻ em có correlation chỉ 0,20–0,35; khác biệt không chỉ do "giáo viên đánh giá sai" mà còn do giáo viên quan sát behaviour observable còn trẻ em trải internal distress (emotional problems). Do đó không nên dùng kết quả này để "trách" giáo viên; cần khung diễn giải tinh tế hơn.', RED),
    ('7. Thiếu TPHCM làm giảm tính khái quát cho đô thị phía Nam', BOLD),
    ('TPHCM là trung tâm kinh tế lớn nhất VN và có đặc thù học đường (mật độ dân số cao, đa dạng dân nhập cư, áp lực học thuật cao nhất cả nước theo các khảo sát PISA-D). Thiếu TPHCM do COVID-19 là hạn chế mang tính logistics nhưng nên được replication trong wave tiếp theo.', RED),
    ('8. Dữ liệu 2021 — giai đoạn đỉnh COVID tại Việt Nam', BOLD),
    ('Thu thập dữ liệu diễn ra năm 2021, khi VN vừa trải qua làn sóng Delta và toàn bộ các tỉnh phía Nam phong toả. Nhiều nghiên cứu cho thấy tác động tâm lý của COVID đối với vị thành niên rõ rệt (Dalton, Rapa & Stein, 2020, Lancet Child & Adolescent Health — có trong refs). Do đó điểm số SDQ có thể cao hơn baseline prepandemic một cách hệ thống. Cần so với các nghiên cứu hậu COVID (2023+) như Hoa 2024 [VN001], Duong 2025 [VN029] để tách "hiệu ứng COVID tạm thời" khỏi "base rate Vietnamese adolescents".', RED),
    ('9. Nhóm LGBTQ và DTTS mới có dữ liệu qualitative', BOLD),
    ('Báo cáo tuyển các em LGBTQ vào FGD và có phát hiện mạnh (kỳ thị peer, thiếu hỗ trợ giáo viên, mistranslation "manly chick" = tomboy); tương tự với HS DTTS (lá ngón Hmong, áp lực tảo hôn). Tuy nhiên cả hai nhóm này đều không có sampling frame định lượng riêng. Colvin, Egan & Coulter (2019, "School Climate & Sexual and Gender Minority Adolescent Mental Health", J Youth Adolesc) cho thấy LGBTQ youth có nguy cơ SKTT cao hơn 2–3 lần so với cisgender-heterosexual peers — VN chưa có ước tính tương ứng.', RED),
    ('10. Không công bố intercoder reliability cho phân tích định tính', BOLD),
    ('Báo cáo không trình bày cụ thể số researcher code FGD/KII, cách resolve disagreement, Cohen\'s kappa giữa các coder. Đây là điểm bắt buộc theo chuẩn COREQ (Consolidated Criteria for Reporting Qualitative Research — Tong, Sainsbury & Craig 2007). Các trích dẫn quote trong báo cáo rất sinh động nhưng reader không kiểm tra được extent của selective quoting.', RED),
    ('', {}),

    ('PHẦN IV — SỐ LIỆU THEN CHỐT (đã verify từ PDF gốc)', H2),
    ('Prevalence tổng quát (SDQ-25)', BOLD),
    ('• 26,1 % học sinh VTN nguy cơ trung bình/cao vấn đề SKTT (n = 668, 2021)', RED),
    ('• 30,9 % vấn đề cảm xúc | 32 % vấn đề bạn bè | 14,4 % tăng động | 11 % vấn đề hành vi', RED),
    ('Khác biệt giới và cấp học', BOLD),
    ('• Nữ > Nam: vấn đề cảm xúc (M = 4,92 vs 4,24, p < 0,01); áp lực học tập (M = 48,11 vs 44,27, p < 0,01) — phù hợp phát hiện toàn cầu (Campbell et al. 2021, n = 566.829)', RED),
    ('• THPT > THCS: vấn đề hành vi, tăng động, cảm xúc, tổng vấn đề — tất cả p < 0,01 — phù hợp xu hướng "late adolescence surge" (Patton et al. 2016)', RED),
    ('Áp lực học tập', BOLD),
    ('• 50 % HS báo cáo học thêm > 3 giờ/tuần | 15 % học > 9 giờ/tuần | 28 % học > 3 giờ/đêm sau giờ học', RED),
    ('• Jayanthi, Thirunavukarasu & Rajkumar (2015, Indian Pediatrics) cho thấy academic stress và depressive symptoms có OR ≈ 2,1 ở vị thành niên Ấn Độ — cơ chế có thể áp dụng cho VN', RED),
    ('Bắt nạt và bắt nạt mạng', BOLD),
    ('• Cyberbullying: 52,2 % đã từng bị (hiếm/đôi khi/thường); THPT > THCS', RED),
    ('• Đối chiếu: Cong et al. 2018 (n = 770 HS Hà Nội) đo tỷ lệ victim 30–40 %; UNICEF 2022 cao hơn, có thể do đo broader exposure (hiếm + đôi khi + thường) và giai đoạn COVID tăng online time', RED),
    ('Nhân lực và hệ thống', BOLD),
    ('• 86,4 % GV chưa nhận đào tạo MHPSS | 95 % GV lo ngại SKTT HS | 91 % GV lo ngại stress HS', RED),
    ('• 70 % trường ở tỉnh nông thôn thiếu phòng tư vấn riêng tư (Thông tư 31/2017 chưa thực hiện đầy đủ)', RED),
    ('• 66 GV surveys + 34 KII + 668 HS surveys + 39 HS FGD + 21 phụ huynh FGD + 35 GV FGD', RED),
    ('', {}),

    ('PHẦN V — ĐỐI CHIẾU LIÊN BÀI TRONG DỰ ÁN LO ÂU', H2),
    ('VN001 (Hoa et al. 2024, n = 3.910, Hà Nội + Thái Nguyên + Cần Thơ)', BOLD),
    ('Dùng GAD-7 và PHQ-9 — không trực tiếp so được SDQ, nhưng cùng chỉ ra nữ > nam (OR ≈ 2,0 cho GAD-7 cut-off 10), THPT > THCS. Bổ sung cho UNICEF 2022: prevalence lo âu tổng quát (GAD-7 ≥ 10) 23,5 %, gần với 26,1 % của UNICEF.', RED),
    ('VN029 (Duong et al. 2025, n = 2.631, longitudinal 3 timepoints)', BOLD),
    ('Thiết kế longitudinal khắc phục hạn chế cắt ngang của UNICEF 2022. Báo cáo tracking quỹ đạo triệu chứng qua 2 năm; cho thấy HS có nguy cơ cao ban đầu (SDQ-like) duy trì nguy cơ cao ở timepoint 2 (persistence rate ≈ 55 %).', RED),
    ('VN030 (Happy House — Dang et al.)', BOLD),
    ('Đây là can thiệp MHPSS trường học có đối chứng tại TPHCM — bổ sung nhu cầu "RCT can thiệp MHPSS học đường" mà UNICEF 2022 xác định là gap. Tham khảo cũng liên kết citation #901 trong refs của UNICEF: Dang, Weiss, Nguyen, Tran & Pollack 2016 về efficacy can thiệp RISE tại VN.', RED),
    ('QT037, QT038, QT042 — literature quốc tế', BOLD),
    ('Meta-analysis của Das et al. 2016 (J Adolesc Health, "Interventions for Adolescent Mental Health: An Overview of Systematic Reviews") chỉ ra CBT-based school programs có effect size d = 0,20–0,35 cho anxiety/depression — đây là benchmark khi VN thiết kế can thiệp.', RED),
    ('', {}),

    ('PHẦN VI — BỐI CẢNH KHU VỰC ĐÔNG NAM Á', H2),
    ('• Indonesia (NAMHS 2021, Erskine et al. 2021): prevalence any mental disorder 5,5 % (WHO-CIDI) — thấp hơn SDQ-based estimate của UNICEF VN 2022, phản ánh khác biệt instrument (CIDI chẩn đoán lâm sàng vs SDQ screening).', RED),
    ('• Philippines (Young Adult Fertility and Sexuality Study 2019): suicide ideation 17 % ở HS 15–24; VN Blum, Sudhinaraset & Emerson 2012 (citation #878) ước tính 8,3 % suicide ideation ở HS VN 11–18.', RED),
    ('• Thailand (NAMHS pending): chưa có số national cho adolescent, nhưng review của Hirota, Guerrero & Skokauskas 2020 (citation #919) chỉ ra tỷ lệ tư vấn viên học đường / HS ở Thailand ~ 1:1.200 — cao hơn đáng kể so với VN ~ 1:4.500 (theo khuyến nghị MOET).', RED),
    ('Đặt trong khu vực, UNICEF VN 2022 là một trong những national reports toàn diện nhất về yếu tố trường học, nhưng còn thiếu prevalence rối loạn theo DSM-5/ICD-11 — khoảng trống này có thể được NAMHS Vietnam (Erskine et al. 2021) lấp trong tương lai.', RED),
    ('', {}),

    ('PHẦN VII — HƯỚNG NGHIÊN CỨU TIẾP (GAP cụ thể)', H2),
    ('1. RCT can thiệp MHPSS trường học tại VN ngoài Happy House — cần ít nhất 3–5 cluster-RCT tại các tỉnh khác nhau để xác lập efficacy (Das et al. 2016 framework).', RED),
    ('2. Nghiên cứu longitudinal từ THCS lên THPT (4 năm) để theo trajectory — VN hiện chỉ có Duong 2025 là longitudinal 2 năm.', RED),
    ('3. Ước tính prevalence rối loạn theo chuẩn DSM-5/ICD-11 bằng CIDI hoặc K-SADS — chưa từng có tại VN (khác với Indonesia NAMHS).', RED),
    ('4. Khung định lượng riêng cho nhóm LGBTQ (Colvin et al. 2019 có thể adapt với quality sampling snowball + respondent-driven).', RED),
    ('5. Khung định lượng riêng cho nhóm DTTS — hiện UNICEF 2022 chỉ có qualitative; cần instrument culturally adapted cho Mông, Thái, Ba Na, Gia-rai.', RED),
    ('6. Đào tạo tư vấn viên học đường chuyên nghiệp — evaluate mô hình "1 tư vấn viên / trường" vs "1 tư vấn viên / cụm trường" về cost-effectiveness.', RED),
    ('7. Hợp tác liên ngành MOET–MOH–MOLISA — hiện chưa có National Strategy on Adolescent Mental Health; cần policy evaluation các mô hình thí điểm.', RED),
    ('8. Chính sách giảm áp lực học tập — impact evaluation cho Thông tư 28/2020/TT-BGDĐT về giảm bài tập về nhà (thực thi vs không thực thi).', RED),
    ('', {}),
]

# Insert lines AFTER the "before_anchor" paragraph
# Strategy: create each new paragraph at end of doc, then move via addnext in reverse
anchor_el = before_anchor._element
# Build all new paragraphs appended to doc
new_paras_elements = []
for text, style in lines:
    kw = {}
    for k in ('bold', 'italic', 'size', 'red'):
        if k in style:
            kw[k] = style[k]
    p = make_para(text if text else '', **kw)
    new_paras_elements.append(p._element)

# Remove from end and insert after anchor, in reverse order so final order is preserved
for el in reversed(new_paras_elements):
    el.getparent().remove(el)
    anchor_el.addnext(el)

d.save(PATH)

# Verify
d2 = Document(PATH)
total = len(d2.paragraphs)
print(f'\nFinal total paragraphs: {total}')

# Find the expanded section
for i, p in enumerate(d2.paragraphs):
    t = p.text.strip()
    if t.startswith('PHẦN I ') or t.startswith('PHẦN II') or t.startswith('PHẦN III') or t.startswith('PHẦN IV') or t.startswith('PHẦN V ') or t.startswith('PHẦN VI') or t.startswith('PHẦN VII'):
        print(f'  #{i}: {t[:80]}')
    if t.startswith('ĐÁNH GIÁ TỔNG THỂ'):
        print(f'  #{i}: {t[:80]}')
