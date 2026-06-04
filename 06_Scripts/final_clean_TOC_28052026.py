# -*- coding: utf-8 -*-
"""Final clean - xoa TOC field rong em them, re-extract page, update bang TOC.
28/05/2026."""
import os, sys, io, shutil, json, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import win32com.client
import pythoncom
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
BACKUP_CLEAN = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026_BEFORE_TOC.docx')

# ============================================================
# STEP 1: Restore from BEFORE_TOC (clean state)
# ============================================================
print("[1] Restoring from BEFORE_TOC...")
shutil.copy(BACKUP_CLEAN, FILE)
print(f"   Restored: {FILE}")


# ============================================================
# STEP 2: Open via Word COM, get heading pages (skip those in tables)
# ============================================================
print("\n[2] Opening via Word COM...")
TEMP = os.path.join(ROOT, 'Luận án TS', '_temp_extract3.docx')
shutil.copy(FILE, TEMP)

pythoncom.CoInitialize()
word = None
headings = []
try:
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = False

    doc = word.Documents.Open(os.path.abspath(TEMP), ReadOnly=False)
    doc.Fields.Update()
    doc.Repaginate()
    print(f"   Pages: {doc.ComputeStatistics(2)}")  # wdStatisticPages=2

    wdActiveEndPageNumber = 3

    for level in [1, 2, 3, 4, 5]:
        for style_name in [f'Heading {level}', f'Tiêu đề {level}']:
            try:
                rng = doc.Content
                find = rng.Find
                find.ClearFormatting()
                find.Style = doc.Styles(style_name)
                find.Text = ''
                find.Forward = True
                find.Wrap = 0
                find.Format = True
                while find.Execute():
                    para = rng.Paragraphs.First
                    text = para.Range.Text.strip().replace('\r', '').replace('\x07', '')
                    if para.Range.Tables.Count > 0:
                        rng.Collapse(0); continue
                    if not text or len(text) > 250:
                        rng.Collapse(0); continue
                    page = para.Range.Information(wdActiveEndPageNumber)
                    headings.append({'level': level, 'text': text, 'page': page})
                    rng.Collapse(0)
            except: pass

    doc.Close(SaveChanges=False)
finally:
    if word is not None:
        try: word.Quit()
        except: pass
    pythoncom.CoUninitialize()
    if os.path.exists(TEMP):
        try: os.remove(TEMP)
        except: pass

# Dedupe
seen = set()
unique = []
for h in headings:
    k = (h['text'], h['page'])
    if k not in seen:
        seen.add(k); unique.append(h)
headings = unique
print(f"   Extracted {len(headings)} unique headings (skipping table-internal)")


# ============================================================
# STEP 3: Match TOC table rows and update
# ============================================================
print("\n[3] Updating TOC table...")
def normalize(s):
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'^[\d\.]+\s*', '', s)
    return s.strip(' .')

lookup = {}
for h in headings:
    k = normalize(h['text'])
    if k and k not in lookup:
        lookup[k] = h['page']
    full = re.sub(r'\s+', ' ', h['text'].strip().lower())
    if full and full not in lookup:
        lookup[full] = h['page']

d = Document(FILE)
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb; break

if toc_table is None:
    print("KHONG TIM THAY BANG TOC")
    sys.exit(1)

updated = 0
not_found = []
all_updates = []
for ri, row in enumerate(toc_table.rows):
    name = row.cells[0].text.strip()
    old_page = row.cells[1].text.strip() if len(row.cells) > 1 else ''
    norm = normalize(name)
    full = re.sub(r'\s+', ' ', name.strip().lower())
    new_page = lookup.get(norm) or lookup.get(full)
    if new_page is None and norm:
        for k, v in lookup.items():
            if k and (len(norm) > 8 and (norm in k or k in norm)):
                new_page = v; break
    if new_page is None:
        not_found.append((ri, name))
        continue
    new_page_str = str(new_page)
    all_updates.append((ri, name, old_page, new_page_str))
    if new_page_str != old_page:
        cell = row.cells[1]
        font_name = 'Times New Roman'
        font_size = Pt(13)
        if cell.paragraphs and cell.paragraphs[0].runs:
            old_r = cell.paragraphs[0].runs[0]
            if old_r.font.name: font_name = old_r.font.name
            if old_r.font.size: font_size = old_r.font.size
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''
        cell.paragraphs[0].text = ''
        new_run = cell.paragraphs[0].add_run(new_page_str)
        new_run.font.name = font_name
        new_run.font.size = font_size
        updated += 1


# Save
d.save(FILE)
print(f"   Updated {updated} page numbers")
print(f"   Not matched: {len(not_found)} rows (giu page cu)")
print(f"\n   Saved: {FILE}")
print(f"   Size: {os.path.getsize(FILE)//1024} KB")

# Show summary
print()
print("=== TOM TAT ===")
print(f"   Total rows TOC: {len(toc_table.rows)}")
print(f"   Matched + updated: {updated}")
print(f"   Matched + same page: {len(all_updates) - updated}")
print(f"   Not matched (heading khong co Heading style): {len(not_found)}")

if not_found:
    print()
    print("Cac muc CHUA co Heading style trong document (page giu nguyen):")
    for ri, name in not_found[:20]:
        print(f"   Row {ri:>2}: {name}")
    if len(not_found) > 20:
        print(f"   ... va {len(not_found)-20} muc khac")
