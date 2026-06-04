# -*- coding: utf-8 -*-
"""GBD 2019 Mental Disorders Collaborators (2022) — V3 chi tiết.
Workflow chuẩn: tiếng Việt thuần + chú thích Anh + đánh dấu trang + bảng Word thật +
hình từ PDF + reference list + phản biện cuối + gap chi tiết.
Lưu ý IP: paraphrase + tóm tắt + trích quote ngắn (KHÔNG reproduce nguyên văn dài).
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\GBD_2019_Mental_Disorders_dich_phan_bien_30042026.docx'
FIG_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\02_Papers-goc\GBD_WHO\figures'

RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)
ORANGE = RGBColor(0xE0, 0x6C, 0x00); RED_BOLD = RGBColor(0xCC, 0, 0)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(9)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def page_marker(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'— — — {text} — — —')
    r.bold=True; r.italic=True; r.font.size=Pt(11)
    r.font.color.rgb=ORANGE; r.font.name='Times New Roman'
def acro(viet, eng, abbrev):
    p=d.add_paragraph()
    r=p.add_run(f'{viet} ({eng} — {abbrev})')
    r.bold=True; r.font.color.rgb=RED_BOLD; r.font.size=Pt(12); r.font.name='Times New Roman'
def crit_para(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run(text); r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
def fig2_ranking_table(title_en, rows_data, age_headers, disorder_colors, vietnamese=False):
    """Build colored ranking table replicating Figure 2 layout.
    rows_data: list of (disorder_name, [r_all, r_0_14, r_15_24, r_25_49, r_50_69, r_70])
    disorder_colors: dict {disorder_name_en: hex_str}
    """
    cap = d.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rcap = cap.add_run(title_en); rcap.bold = True; rcap.italic = True
    rcap.font.size = Pt(11); rcap.font.color.rgb = ORANGE; rcap.font.name = 'Times New Roman'
    cols = ['Disorder' if not vietnamese else 'Rối loạn'] + age_headers
    n_cols = len(cols)
    t = d.add_table(rows=1+len(rows_data), cols=n_cols)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    widths = [4.5, 1.6, 1.6, 1.6, 1.6, 1.6, 1.4]
    for row in t.rows:
        for ci in range(n_cols): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(cols):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name='Times New Roman'; r.font.size = Pt(9)
        shade(c, '2E5A88')
        # White text for header
        for p in c.paragraphs:
            for r in p.runs: r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    for ri, (disorder_name, ranks) in enumerate(rows_data):
        # First cell: disorder name with color
        key = disorder_name if not vietnamese else rows_data[ri][0]
        # Lookup color by index (since vietnamese names differ)
        color_hex = disorder_colors[ri]
        c = t.rows[ri+1].cells[0]; c.text = disorder_name
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name='Times New Roman'; r.font.size = Pt(9); r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        shade(c, color_hex)
        # Other cells: rank with light tint
        for ci, val in enumerate(ranks):
            cc = t.rows[ri+1].cells[ci+1]; cc.text = str(val)
            for p in cc.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in p.runs: r.font.name='Times New Roman'; r.font.size = Pt(10); r.bold = True
            # Light shade if NA
            if str(val) == 'NA':
                shade(cc, 'D0D0D0')
            else:
                # Light tint of the disorder color
                shade(cc, color_hex + '40' if len(color_hex)==6 else color_hex)
                # Better: use a fixed light shade
                shade(cc, lighten_hex(color_hex))
    d.add_paragraph()


def lighten_hex(hex6):
    """Return a lighter version of hex color for cell tint."""
    r = int(hex6[0:2], 16); g = int(hex6[2:4], 16); b = int(hex6[4:6], 16)
    r = min(255, r + (255-r)*7//10); g = min(255, g + (255-g)*7//10); b = min(255, b + (255-b)*7//10)
    return f'{r:02X}{g:02X}{b:02X}'


def insert_image(fname, caption, width_cm=15):
    path = os.path.join(FIG_DIR, fname)
    if not os.path.exists(path):
        nr(f'[Hình không tải được: {fname}]', italic=True, color=GRAY, size=11); return
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.add_picture(path, width=Cm(width_cm))
    cap = d.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rcap = cap.add_run(caption)
    rcap.italic = True; rcap.bold = True; rcap.font.size = Pt(10)
    rcap.font.color.rgb = ORANGE; rcap.font.name = 'Times New Roman'

# ========== TRANG BÌA ==========
title("BẢN DỊCH + PHẢN BIỆN CHI TIẾT", 14)
title("GÁNH NẶNG TOÀN CẦU, KHU VỰC VÀ QUỐC GIA", 14)
title("CỦA 12 RỐI LOẠN TÂM THẦN Ở 204 QUỐC GIA", 14)
title("VÀ VÙNG LÃNH THỔ, 1990–2019", 14)
nr("")
subtitle("Global, regional, and national burden of 12 mental disorders in 204 countries "
         "and territories, 1990–2019: a systematic analysis for the Global Burden of "
         "Disease Study 2019")
nr("")
subtitle("GBD 2019 Mental Disorders Collaborators (~150 tác giả)")
subtitle("The Lancet Psychiatry, Tập 9(2): 137–150 — đăng 11/01/2022")
subtitle("DOI: 10.1016/S2215-0366(21)00395-3 — PMID 35026139")
nr("")
subtitle("Trợ lý nghiên cứu — 30/04/2026 — tiếng Việt thuần — chú thích Anh trong ngoặc")
subtitle("PDF gốc trong dự án: 02_Papers-goc/GBD_WHO/GBD_2019_Mental_Disorders_Lancet_Psychiatry_2022.pdf "
         "(1,1 MB / 14 trang / mở Open Access CC BY 4.0)")

# ========== THÔNG TIN THƯ MỤC ==========
H1("THÔNG TIN THƯ MỤC")
tbl(['Mục', 'Nội dung'], [
    ['Tên bài (tiếng Anh)', 'Global, regional, and national burden of 12 mental disorders in 204 countries '
     'and territories, 1990–2019: a systematic analysis for the Global Burden of Disease Study 2019'],
    ['Tạm dịch tên bài', 'Gánh nặng toàn cầu, khu vực và quốc gia của 12 rối loạn tâm thần ở 204 quốc gia '
     'và vùng lãnh thổ, giai đoạn 1990–2019: phân tích hệ thống cho Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019'],
    ['Tác giả', 'GBD 2019 Mental Disorders Collaborators (~150 tác giả) — '
     'liên hệ tác giả chính: Dr Alize Ferrari (a.ferrari@uq.edu.au) — '
     'Queensland Centre for Mental Health Research, University of Queensland, Brisbane QLD 4108, Úc'],
    ['Tạp chí', 'The Lancet Psychiatry'],
    ['Tập / Số / Trang', 'Tập 9(2): trang 137–150 (số tháng 2/2022)'],
    ['Năm + ngày đăng online', '11/01/2022'],
    ['DOI', '10.1016/S2215-0366(21)00395-3'],
    ['PMID', '35026139'],
    ['Loại bài', 'Phân tích hệ thống dữ liệu dịch tễ học toàn cầu 30 năm (1990–2019)'],
    ['Quy mô', '12 rối loạn tâm thần × 23 nhóm tuổi × 2 giới × 204 quốc gia/lãnh thổ × 30 năm'],
    ['Thiết kế', 'Phân tích thứ cấp dữ liệu dịch tễ học có hệ thống — '
     'phân tích hồi quy meta-Bayesian dùng công cụ DisMod-MR 2.1'],
    ['Tài trợ chính', 'Bill & Melinda Gates Foundation; Australian National Health and Medical '
     'Research Council; Queensland Department of Health, Úc'],
    ['Xung đột lợi ích', 'Một số tác giả công bố nhận grant/thù lao từ MQ Mental Health, '
     'NHMRC, NIH, Janssen, Lundbeck, Sanofi, Vistagen, Boehringer-Ingelheim, Roche/Genentech '
     '— ngoài bài này; phần lớn tác giả không có xung đột'],
    ['Trạng thái mở', 'Open Access dưới giấy phép CC BY 4.0 (free để đọc + chia sẻ + thích nghi với ghi nguồn)'],
    ['Đường dẫn truy cập', 'https://www.thelancet.com/journals/lanpsy/article/PIIS2215-0366(21)00395-3/fulltext'],
    ['Mã trong kho dự án', 'Đề xuất: QT_GBD2019_MD'],
], [4.0, 12.0])

# ========== BẢNG TỪ VIẾT TẮT ==========
H1("BẢNG TỪ VIẾT TẮT DÙNG TRONG TÀI LIỆU NÀY")
nr("Lần đầu mỗi từ viết tắt xuất hiện trong văn bản dưới đây sẽ được tô MÀU ĐỎ ĐẬM "
   "+ chú thích inline.")
tbl(['Tiếng Việt', 'Tiếng Anh đầy đủ', 'Viết tắt'], [
    ['Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2019', 'Global Burden of Disease Study 2019', 'GBD 2019'],
    ['Năm sống điều chỉnh theo khuyết tật', 'Disability-Adjusted Life-Year', 'DALY'],
    ['Năm sống với khuyết tật', 'Years Lived with Disability', 'YLD'],
    ['Năm sống bị mất do tử vong sớm', 'Years of Life Lost (to premature mortality)', 'YLL'],
    ['Khoảng tin cậy không chắc chắn 95%', '95% Uncertainty Interval', '95% UI'],
    ['Trọng số khuyết tật', 'Disability weight', '—'],
    ['Mô hình hồi quy meta-Bayesian DisMod', 'Bayesian meta-regression tool DisMod-MR 2.1', 'DisMod-MR 2.1'],
    ['Hệ thống Dữ liệu Y tế Toàn cầu', 'Global Health Data Exchange', 'GHDx'],
    ['Sổ tay chẩn đoán + thống kê rối loạn tâm thần phiên bản 4 + 5',
     'Diagnostic and Statistical Manual of Mental Disorders, 4th + 5th edition', 'DSM-IV / DSM-5'],
    ['Phân loại bệnh quốc tế phiên bản 10 + 11',
     'International Classification of Diseases, 10th + 11th revision', 'ICD-10 / ICD-11'],
    ['Rối loạn tăng động giảm chú ý',
     'Attention-Deficit Hyperactivity Disorder', 'ADHD'],
    ['Khuyết tật trí tuệ phát triển không rõ nguyên nhân',
     'Idiopathic Developmental Intellectual Disability', '—'],
    ['Rối loạn phổ tự kỷ', 'Autism Spectrum Disorder', '—'],
    ['Rối loạn ăn uống (chán ăn tâm thần + cuồng ăn tâm thần)',
     'Eating Disorders (anorexia nervosa + bulimia nervosa)', '—'],
    ['Rối loạn lo âu (gộp tất cả phân loại con)',
     'Anxiety Disorders (combined estimate of all subtypes)', '—'],
    ['Rối loạn trầm cảm (trầm cảm chính + dysthymia)',
     'Depressive Disorders (Major Depressive Disorder + Dysthymia)', 'MDD + Dysthymia'],
    ['Rối loạn lưỡng cực', 'Bipolar Disorder', '—'],
    ['Tâm thần phân liệt', 'Schizophrenia', '—'],
    ['Rối loạn hành vi', 'Conduct Disorder', '—'],
], [5.0, 8.5, 2.5])

# ========== PHẦN 1 — TÓM TẮT (paraphrase) ==========
H1("PHẦN 1 — TÓM TẮT NGHIÊN CỨU (Summary, Abstract)")
page_marker('Trang 137, Lancet Psychiatry, Tập 9(2), 02/2022')

H2("1.1. Bối cảnh (Background)")
nr("Các rối loạn tâm thần được đưa vào phân tích GBD 2019 gồm 12 nhóm: rối loạn trầm "
   "cảm, rối loạn lo âu, rối loạn lưỡng cực, tâm thần phân liệt, rối loạn phổ tự kỷ, "
   "rối loạn hành vi, rối loạn tăng động giảm chú ý (ADHD), rối loạn ăn uống, khuyết "
   "tật trí tuệ phát triển không rõ nguyên nhân, và một nhóm \"các rối loạn tâm thần "
   "khác\" còn lại. Mục tiêu nghiên cứu: ước tính tỉ lệ hiện mắc, năm sống điều chỉnh "
   "theo khuyết tật (DALY), năm sống với khuyết tật (YLD), năm sống bị mất do tử vong "
   "sớm (YLL) cho các rối loạn tâm thần ở cấp toàn cầu, khu vực, quốc gia, từ 1990 "
   "đến 2019.")

H2("1.2. Phương pháp (Methods)")
nr("Nhóm tác giả ước tính tỉ lệ hiện mắc + gánh nặng từ GBD 2019 cho 12 rối loạn, "
   "phân theo 23 nhóm tuổi, nam/nữ, 204 quốc gia và vùng lãnh thổ, giai đoạn 1990–2019. "
   "DALYs được tính bằng tổng YLD + YLL. Quy trình tổng quan hệ thống dữ liệu trên 4 "
   "cơ sở: PsycINFO, Embase, PubMed, Global Health Data Exchange (GHDx). Sau đó tỉ lệ "
   "hiện mắc được ước tính bằng phân tích hồi quy meta-Bayesian (DisMod-MR 2.1) — "
   "phân theo rối loạn, tuổi, giới, năm và địa điểm. YLD được tính bằng tỉ lệ hiện mắc "
   "× trọng số khuyết tật. Tử vong + YLL CHỈ được tính cho rối loạn ăn uống (chán ăn "
   "tâm thần + cuồng ăn tâm thần) — đây là 2 rối loạn duy nhất được nhận diện làm "
   "nguyên nhân tử vong trực tiếp trong khung GBD 2019.")

H2("1.3. Phát hiện chính (Findings)")
nr("• 1990–2019: tổng số DALY toàn cầu do rối loạn tâm thần TĂNG TỪ 80,8 triệu (95% "
   "UI 59,5–105,9) lên 125,3 triệu (95% UI 93,0–163,2) — tăng tổng 55,1%.")
nr("• Tỉ trọng DALY toàn cầu do rối loạn tâm thần tăng từ 3,1% (95% UI 2,4–3,9) lên "
   "4,9% (95% UI 3,9–6,1) — tức gần GẤP RƯỠI sau 30 năm.")
nr("• Tỉ lệ DALY chuẩn hoá theo tuổi GẦN NHƯ KHÔNG ĐỔI: 1581,2 (95% UI 1170,9–2061,4) "
   "trên 100.000 dân năm 1990 vs 1566,2 (95% UI 1160,1–2042,8) năm 2019.")
nr("• Năm 2019: YLD chiếm phần lớn gánh nặng — 125,3 triệu YLD (95% UI 93,0–163,2; "
   "tương đương 14,6% [12,2–16,8] tổng YLD toàn cầu) do rối loạn tâm thần.")
nr("• Rối loạn ăn uống ghi nhận 17.361,5 YLL (95% UI 15.518,5–21.459,8) — DUY NHẤT "
   "trong 12 rối loạn được tính tử vong trực tiếp.")
nr("• Tỉ lệ DALY chuẩn hoá toàn cầu: NAM = 1.426,5 (1.056,4–1.869,5) trên 100.000; "
   "NỮ = 1.703,3 (1.261,5–2.237,8) trên 100.000. → NỮ cao hơn NAM ~19%.")
nr("• Tỉ lệ DALY chuẩn hoá CAO NHẤT ở 3 vùng: Australasia, Mỹ Latin nhiệt đới (Tropical "
   "Latin America), Bắc Mỹ thu nhập cao (high-income North America).")

H2("1.4. Diễn giải (Interpretation)")
nr("GBD 2019 cho thấy rối loạn tâm thần VẪN nằm trong TOP 10 nguyên nhân gây gánh "
   "nặng toàn cầu, KHÔNG có bằng chứng giảm gánh nặng kể từ 1990. Số YLL ước tính "
   "cho rối loạn tâm thần là CỰC THẤP và KHÔNG phản ánh tử vong sớm thực tế ở người "
   "có rối loạn tâm thần. Tác giả khuyến nghị: nghiên cứu thêm về đường dẫn nhân quả "
   "giữa rối loạn tâm thần và các kết cục y tế gây tử vong khác — để có thể đưa vào "
   "khung GBD. Ngoài ra: cần CAN THIỆP PHÒNG NGỪA + ĐIỀU TRỊ HIỆU QUẢ + có sự PHỐI "
   "HỢP của chính phủ + cộng đồng y tế toàn cầu để giảm gánh nặng này.")

# ========== PHẦN 2 — PHƯƠNG PHÁP ==========
H1("PHẦN 2 — PHƯƠNG PHÁP CHI TIẾT")
page_marker('Trang 138-141, Lancet Psychiatry, Tập 9(2), 02/2022')

H2("2.1. Định nghĩa case (Case definitions)")
nr("GBD 2019 bao gồm 369 bệnh và chấn thương được tổ chức theo cấp bậc 4 mức nguyên "
   "nhân. Trong nghiên cứu này, 12 rối loạn tâm thần thuộc các mức của hệ thống GBD. "
   "Định nghĩa case TUÂN THEO tiêu chí DSM-IV-TR hoặc ICD-10 (hai hệ thống chuẩn được "
   "đa số khảo sát SKTT toàn cầu sử dụng). Lý do: cần đảm bảo tính so sánh giữa các "
   "khảo sát ở các quốc gia khác nhau.")
nr("Lưu ý quan trọng: ĐỂ TRÁNH THIÊN LỆCH NHỚ LẠI (recall bias) đối với nhiều rối loạn "
   "trong các thước đo \"tỉ lệ trong cuộc đời\" (lifetime prevalence), nghiên cứu chỉ "
   "BAO GỒM các báo cáo tỉ lệ \"trong năm qua\" (past-year) hoặc \"thấp hơn\" cho TẤT "
   "CẢ rối loạn — TRỪ rối loạn lưỡng cực + rối loạn phổ tự kỷ (vẫn dùng thiết kế dài "
   "hạn — prospective design).")

H2("2.2. Ước tính YLD (Years Lived with Disability)")
nr("YLD được tính bằng cách nhân tỉ lệ hiện mắc ở các mức độ nặng khác nhau với trọng "
   "số khuyết tật (disability weight) tương ứng. Trọng số khuyết tật định lượng mức "
   "độ MẤT MÁT SỨC KHỎE liên quan đến từng di chứng (sequela) hoặc hậu quả của bệnh "
   "tật/chấn thương — phổ điểm 0 (sức khỏe hoàn hảo) đến 1 (tử vong).")

H2("2.3. Nguồn dữ liệu (Data sources)")
nr("Để tổng hợp dữ liệu dịch tễ học cần thiết cho ước tính YLD, nhóm nghiên cứu thực "
   "hiện tổng quan tài liệu hệ thống trên: PsycINFO, Embase, PubMed + nguồn xám "
   "(grey literature) + tham vấn chuyên gia. KHÔNG giới hạn ngôn ngữ. Cũng dùng dữ "
   "liệu trong Global Health Data Exchange (GHDx — http://ghdx.healthdata.org/) — "
   "kho khảo sát đa quốc gia + nguồn dữ liệu được cộng tác viên GBD đề xuất.")
nr("Tiêu chí lựa chọn nghiên cứu:")
nr("• Khảo sát đăng từ 01/01/1980 đến nay")
nr("• Có lấy mẫu xác suất (probability sampling) đại diện cho dân số chung")
nr("• Loại trừ: khảo sát chỉ trên dân số phụ (vd: chỉ tù nhân)")
nr("• Loại trừ: mẫu lâm sàng (treatment samples), TRỪ trường hợp có thể nắm bắt mọi "
   "ca bệnh trong dân số (vd: rối loạn phổ tự kỷ với hiệu chỉnh thiên lệch)")

H2("2.4. Mô hình dịch tễ học (Epidemiological disease models)")
nr("Dữ liệu dịch tễ học đã được phân tích theo 2 bước:")
nr("• Bước 1: Kiểm tra + điều chỉnh các thiên lệch trong ước tính dịch tễ giữa các "
   "nghiên cứu")
nr("• Bước 2: Ước tính \"chuẩn vàng\" (gold-standard) — dùng phương pháp thu thập dữ "
   "liệu mong muốn + KHÔNG cần điều chỉnh thiên lệch — được mô hình hoá trong phân "
   "tích hồi quy meta-regression")
nr("Đối với mỗi rối loạn, nhóm xác định các nguồn thiên lệch CHÍNH trong dữ liệu trích "
   "xuất + các nguồn sai số đo lường: loại nhớ lại (point/12-month/lifetime), công cụ "
   "khảo sát (chẩn đoán hay thang triệu chứng), người phỏng vấn (lay/clinician). "
   "Ước tính có thiên lệch được coi là ước tính thay thế (alternative estimates) cho "
   "ước tính chuẩn vàng + được điều chỉnh.")
nr("Mô hình DisMod-MR 2.1 (Bayesian meta-regression tool) được dùng để gộp dữ liệu từ "
   "nhiều nguồn khác nhau → tạo ra các ước tính nhất quán nội tại về tỉ lệ hiện mắc, "
   "tỉ lệ mới mắc, tỉ lệ thuyên giảm, và tỉ lệ tử vong dư thừa — phân theo tuổi, giới, "
   "địa điểm, năm. DisMod-MR 2.1 dùng các đồng biến cấp địa điểm (location-level "
   "covariates) để dự đoán tỉ lệ ở các vị trí thiếu dữ liệu.")
nr("3 đồng biến chính cho rối loạn trầm cảm chính + lo âu:")
nr("(1) Tỉ lệ tử vong trung bình 10 năm trước do chiến tranh và khủng bố (đại diện "
   "cho mức độ xung đột)")
nr("(2) Chỉ số trải nghiệm tiêu cực Gallup (Gallup Negative Experience Index) — đo "
   "trải nghiệm hằng ngày: đau thể chất, lo lắng, buồn, căng thẳng, tức giận")
nr("(3) Tỉ lệ trầm cảm do 2 yếu tố nguy cơ: bạo lực bạn tình + lạm dụng tình dục thời "
   "thơ ấu")

H2("2.5. Tỉ lệ mức độ nặng + Trọng số khuyết tật")
nr("Tỉ lệ phân bố mức độ nặng (severity proportions) — vd: nhẹ / trung bình / nặng — "
   "được áp dụng vào tổng tỉ lệ hiện mắc đã ước tính bởi DisMod-MR 2.1 → tỉ lệ ở "
   "mỗi mức độ.")
nr("Trọng số khuyết tật được tính từ các khảo sát cộng đồng ở Bangladesh, Indonesia, "
   "Peru, Tanzania, Hoa Kỳ, Hungary, Italia, Thuỵ Điển, Hà Lan + 1 khảo sát trực "
   "tuyến mở (English, Spanish, Mandarin). Phổ điểm 0 (sức khỏe hoàn hảo) đến 1 (tử "
   "vong).")

H2("2.6. Điều chỉnh đồng mắc (Comorbidity adjustments)")
nr("Vì gánh nặng do mỗi nguyên nhân GBD được ước tính RIÊNG, nhóm dùng phương pháp "
   "MÔ PHỎNG (simulation method) để điều chỉnh đồng mắc. Mô phỏng dân số 40.000 cá "
   "thể theo địa điểm, tuổi, giới, năm — mỗi cá thể được tiếp xúc với xác suất độc "
   "lập có sự kết hợp di chứng trong GBD 2019. Điều chỉnh đồng mắc tính trung bình "
   "khác biệt giữa trọng số khuyết tật của cá thể có MỘT di chứng so với người có "
   "NHIỀU di chứng.")

H2("2.7. Ước tính YLL (Years of Life Lost) + DALY")
nr("YLL = số ca tử vong cụ thể × số năm sống mong đợi còn lại tại thời điểm chết "
   "(theo bảng tuổi thọ tham chiếu). Trong khuôn khổ GBD, MỖI ca tử vong CHỈ được gán "
   "cho 1 nguyên nhân nền tảng theo phân loại ICD. Tử vong + YLL CHỈ được tính cho "
   "rối loạn ăn uống — vì đó là 2 rối loạn tâm thần duy nhất được nhận diện làm "
   "nguyên nhân tử vong nền tảng trong GBD 2019.")
nr("LƯU Ý QUAN TRỌNG: Tử vong sớm ở người có rối loạn tâm thần do nguyên nhân khác "
   "(vd: tự tử, bệnh tim mạch, bệnh hô hấp, neoplasm, tiểu đường, viêm) — KHÔNG được "
   "gán cho rối loạn tâm thần trong khuôn khổ GBD → YLL ước tính cho rối loạn tâm "
   "thần KHÔNG phản ánh tử vong sớm thực tế.", color=BLUE, bold=True)
nr("DALY = YLD + YLL. Khoảng tin cậy không chắc chắn 95% (95% UI) được ước tính từ "
   "phân vị 25 và 975 của 1.000 lần phân phối hậu nghiệm tại mỗi bước.")

H2("2.8. Phân cấp địa lý")
nr("GBD 2019 phân chia thành 7 super-region + 21 region + 204 quốc gia/lãnh thổ. "
   "Tuổi: 23 nhóm (từ <1 tuổi đến ≥95 tuổi). Mỗi năm từ 1990 đến 2019.")
nr("95% UI (khoảng không chắc chắn) cho tất cả ước tính được tính từ phân vị 25 và "
   "975 của 1.000 lần phân phối hậu nghiệm. Excel + maptools R (phiên bản 3.6.3) "
   "được dùng để tạo bảng + hình.")

# ========== PHẦN 3 — KẾT QUẢ ==========
H1("PHẦN 3 — KẾT QUẢ CHI TIẾT")
page_marker('Trang 141-145, Lancet Psychiatry, Tập 9(2), 02/2022')

H2("3.1. Tổng quan kết quả")
nr("Năm 2019: ước tính 970,1 triệu ca rối loạn tâm thần toàn cầu (95% UI 900,9–1.044,4) "
   "— TĂNG 48,1% từ 654,8 triệu (95% UI 603,6–708,1) năm 1990. KHÔNG có thay đổi rõ "
   "rệt ở tỉ lệ chuẩn hoá theo tuổi → tăng tổng số ca chủ yếu DO TĂNG DÂN SỐ + DÂN SỐ "
   "GIÀ HOÁ.")

H2("3.2. Tỉ lệ chuẩn hoá theo tuổi của 12 rối loạn (Bảng 1 tái tạo)")
nr("Tỉ lệ chuẩn hoá theo tuổi (age-standardised prevalence) trên 100.000 dân, năm 1990 "
   "vs năm 2019. Em tái tạo các số liệu chính:")
tbl(['Rối loạn', 'Tổng 1990 (95% UI)', 'Tổng 2019 (95% UI)', 'Nam 2019', 'Nữ 2019'], [
    ['TỔNG 12 RỐI LOẠN', '12.579,3 (11.634,4–13.552,2)',
     '12.262,0 (11.382,9–13.213,3)', '11.727,3 (10.835,7–12.693,9)', '12.760,0 (11.831,7–13.763,1)'],
    ['Rối loạn lo âu', '3.791,6 (3.194,0–4.476,6)', '3.779,5 (3.181,1–4.473,3)',
     '2.859,8 (2.397,0–3.379,9)', '4.694,7 (3.945,6–5.576,9)'],
    ['Rối loạn trầm cảm', '3.486,2 (3.140,8–3.855,7)', '3.440,1 (3.097,0–3.817,6)',
     '2.713,3 (2.438,3–3.013,1)', '4.158,4 (3.746,9–4.616,3)'],
    ['Khuyết tật trí tuệ phát triển', '1.641,9 (1.028,3–2.278,2)', '1.426,5 (873,6–1.991,7)',
     '1.436,4 (860,4–2.027,8)', '1.415,4 (891,3–1.954,5)'],
    ['Rối loạn tăng động giảm chú ý (ADHD)', '1.240,5 (909,6–1.647,1)', '1.131,9 (831,7–1.494,5)',
     '1.611,6 (1.184,8–2.134,1)', '631,0 (455,7–846,5)'],
    ['Rối loạn hành vi', '537,9 (388,2–699,0)', '559,0 (405,0–722,3)',
     '711,2 (530,5–904,0)', '397,3 (263,8–545,5)'],
    ['Rối loạn lưỡng cực', '490,1 (411,0–576,5)', '489,8 (407,5–580,6)',
     '466,9 (388,5–552,9)', '512,8 (425,6–609,0)'],
    ['Rối loạn phổ tự kỷ', '372,8 (309,1–444,9)', '369,4 (305,9–441,2)',
     '560,1 (465,2–667,3)', '176,3 (143,0–214,5)'],
    ['Tâm thần phân liệt', '289,9 (249,8–333,2)', '287,4 (246,2–330,9)',
     '302,7 (259,7–348,4)', '272,0 (232,7–313,7)'],
    ['Rối loạn ăn uống', '150,5 (113,1–192,1)', '174,0 (130,1–222,1)',
     '117,9 (84,6–156,1)', '231,5 (175,1–291,4)'],
], [3.5, 3.0, 3.0, 2.5, 2.5])
nr("Quan sát: rối loạn lo âu và trầm cảm là HAI rối loạn phổ biến nhất xuyên 30 năm. "
   "Nữ cao hơn nam ở: trầm cảm, lo âu, ăn uống. Nam cao hơn nữ ở: ADHD + tự kỷ + "
   "hành vi + tâm thần phân liệt nhẹ.", bold=True)

H2("3.3. Tỉ lệ DALY chuẩn hoá theo tuổi cho rối loạn tâm thần (toàn cầu)")
nr("Tỉ lệ DALY chuẩn hoá theo tuổi: 1.426,5 (95% UI 1.056,4–1.869,5) trên 100.000 "
   "dân ở NAM; 1.703,3 (95% UI 1.261,5–2.237,8) trên 100.000 dân ở NỮ.")
nr("Tỉ lệ DALY toàn cầu cho TỔNG rối loạn tâm thần: 1.566,2 (95% UI 1.160,1–2.042,8) "
   "trên 100.000 dân năm 2019. Số tổng DALY: 125,3 triệu (95% UI 93,0–163,2).")
nr("Đóng góp của từng rối loạn vào tổng gánh nặng SKTT toàn cầu năm 2019:")
nr("• Trầm cảm: 37,3% (95% UI 32,3–43,0)")
nr("• Lo âu: 22,9% (95% UI 18,6–27,5)")
nr("• Tâm thần phân liệt: 12,2% (95% UI 9,6–15,2)")
nr("• Các rối loạn khác chiếm phần còn lại")

insert_image('GBD2019_Figure1_DALYs_pyramid.png',
             'Hình 1 (trang 145 PDF gốc) — Phân bố DALY toàn cầu theo rối loạn tâm thần, '
             'giới tính và nhóm tuổi, năm 2019. Pyramid hai bên cho thấy NAM (trái) và '
             'NỮ (phải); màu khác nhau cho từng rối loạn (đỏ=trầm cảm, vàng=lo âu, '
             'xanh dương=tâm thần phân liệt, v.v.)',
             width_cm=15)

H2("3.4. Phân bố theo nhóm tuổi (chi tiết)")
nr("Một số phát hiện QUAN TRỌNG về phân bố theo tuổi:")
nr("• 80,6% gánh nặng do rối loạn tâm thần xảy ra ở tuổi LAO ĐỘNG (16-65 tuổi)")
nr("• 9,2% gánh nặng còn lại xảy ra ở người TRẺ HƠN 16 tuổi")
nr("• 23,2% gánh nặng do rối loạn tâm thần ở trẻ em + vị thành niên TOÀN CẦU nằm ở "
   "khu vực TIỂU SAHARA (sub-Saharan Africa)")
nr("• Gánh nặng XUẤT HIỆN trước 5 tuổi ở người có khuyết tật trí tuệ phát triển + "
   "rối loạn phổ tự kỷ")
nr("• Tiếp tục rõ ở tuổi LỚN HƠN ở người có rối loạn trầm cảm, lo âu và tâm thần "
   "phân liệt")
nr("• Số DALY tăng đều đặn trong thời thơ ấu + vị thành niên, ĐẠT ĐỈNH ở 25-34 tuổi, "
   "rồi giảm đều sau 35 tuổi")

H2("3.5. Xếp hạng YLD và DALY của rối loạn tâm thần (Hình 2)")
nr("Trong tổng số 25 nguyên nhân gây gánh nặng hàng đầu năm 2019:")
nr("• YLD: trầm cảm xếp HẠNG 2; lo âu xếp HẠNG 8; tâm thần phân liệt HẠNG 20")
nr("• DALY: trầm cảm xếp HẠNG 13; lo âu HẠNG 24; tâm thần phân liệt HẠNG 42")
nr("• Rối loạn tâm thần là nguyên nhân thứ 13 GÂY DALY toàn cầu năm 1990 và nguyên "
   "nhân thứ 7 GÂY YLD năm 2019")
nr("• Rối loạn tâm thần là nguyên nhân thứ 2 GÂY YLD toàn cầu cả năm 1990 + 2019")
nr("• Trong nhóm rối loạn tâm thần, trầm cảm xếp CAO NHẤT ở MỌI nhóm tuổi TRỪ "
   "0-14 tuổi (thay vào đó là rối loạn HÀNH VI xếp cao nhất ở nhóm này)")

# ===== HÌNH 2 — TÁI TẠO DƯỚI DẠNG BẢNG CÓ MÀU (2 bảng: English + Việt) =====
nr("Hình 2 trong PDF gốc thực chất là MỘT BẢNG dữ liệu xếp hạng — em tái tạo đầy đủ "
   "dưới dạng bảng Word có màu (mỗi hàng = 1 rối loạn với màu riêng theo PDF gốc, ô NA "
   "màu xám). Bảng dưới: trên là TIẾNG ANH (theo PDF gốc), dưới là DỊCH TIẾNG VIỆT.")
nr("Cách đọc: số trong ô = THỨ HẠNG của rối loạn đó trong tổng các nguyên nhân Level 3 "
   "của GBD 2019. SỐ CÀNG NHỎ = gánh nặng càng CAO. Ô NA = không ước tính cho nhóm tuổi.")

DISORDER_COLORS = [
    '8B1A1A',  # Depressive: dark red/burgundy
    'E67E22',  # Anxiety: orange
    '6A0DAD',  # Schizophrenia: dark purple
    '808000',  # Other mental disorders: olive
    '2E8B57',  # Bipolar: sea green
    'C71585',  # Conduct: pink/magenta
    '8B4513',  # Idiopathic intellectual: brown
    '008B8B',  # Autism: teal/cyan
    '696969',  # Eating: dim gray
    '4682B4',  # ADHD: steel blue
]

yld_rows_en = [
    ('Depressive disorders',          [2,  24, 2,  3,   5,   11]),
    ('Anxiety disorders',             [8,  8,  4,  6,   16,  19]),
    ('Schizophrenia',                 [20, 92, 22, 9,   19,  36]),
    ('Other mental disorders',        [27, 94, 36, 19,  22,  27]),
    ('Bipolar disorder',              [28, 54, 12, 20,  27,  45]),
    ('Conduct disorder',              [38, 5,  13, 'NA','NA','NA']),
    ('Idiopathic dev. intel. disab.', [43, 18, 30, 44,  64,  87]),
    ('Autism spectrum disorders',     [46, 23, 32, 42,  52,  63]),
    ('Eating disorders',              [55, 65, 28, 36, 'NA','NA']),
    ('ADHD',                          [84, 39, 60, 86,  133, 152]),
]
VI_NAMES = [
    'Rối loạn trầm cảm', 'Rối loạn lo âu', 'Tâm thần phân liệt',
    'Các rối loạn tâm thần khác', 'Rối loạn lưỡng cực', 'Rối loạn hành vi',
    'Khuyết tật trí tuệ phát triển', 'Rối loạn phổ tự kỷ',
    'Rối loạn ăn uống', 'Rối loạn tăng động giảm chú ý (ADHD)',
]
yld_rows_vi = [(VI_NAMES[i], yld_rows_en[i][1]) for i in range(10)]

daly_rows_en = [
    ('Depressive disorders',          [13,  57,  4,  6,   13,  28]),
    ('Anxiety disorders',             [24,  25,  7,  15,  33,  43]),
    ('Schizophrenia',                 [42,  125, 42, 22,  41,  82]),
    ('Other mental disorders',        [64,  127, 59, 36,  55,  66]),
    ('Bipolar disorder',              [67,  98,  32, 39,  62,  94]),
    ('Conduct disorder',              [84,  22,  34, 'NA','NA','NA']),
    ('Idiopathic dev. intel. disab.', [90,  49,  54, 77,  122, 132]),
    ('Autism spectrum disorders',     [92,  56,  56, 73,  104, 120]),
    ('Eating disorders',              [110, 105, 51, 65, 'NA','NA']),
    ('ADHD',                          [145, 84,  87, 135, 154, 159]),
]
daly_rows_vi = [(VI_NAMES[i], daly_rows_en[i][1]) for i in range(10)]

age_hdr_en = ['All ages', '0-14', '15-24', '25-49', '50-69', '>=70']
age_hdr_vi = ['Mọi tuổi', '0-14', '15-24', '25-49', '50-69', '>=70']

fig2_ranking_table('YLDs — Ranking of mental disorders by age group, 2019 (Figure 2 PDF page 146 — English)',
    yld_rows_en, age_hdr_en, DISORDER_COLORS)
fig2_ranking_table('YLDs — Xếp hạng rối loạn tâm thần theo nhóm tuổi, 2019 (Hình 2 PDF trang 146 — Tiếng Việt)',
    yld_rows_vi, age_hdr_vi, DISORDER_COLORS, vietnamese=True)
fig2_ranking_table('DALYs — Ranking of mental disorders by age group, 2019 (Figure 2 PDF page 146 — English)',
    daly_rows_en, age_hdr_en, DISORDER_COLORS)
fig2_ranking_table('DALYs — Xếp hạng rối loạn tâm thần theo nhóm tuổi, 2019 (Hình 2 PDF trang 146 — Tiếng Việt)',
    daly_rows_vi, age_hdr_vi, DISORDER_COLORS, vietnamese=True)

nr("Chú thích Hình 2 (theo PDF gốc): Mental disorders were ranked out of all Level 3 "
   "causes within the Global Burden of Diseases, Injuries, and Risk Factors Study. "
   "Disorders are ordered from highest to lowest ranking for the all ages group. Each "
   "colour represents a different mental disorder. Grey cells marked NA show disorders "
   "for which burden was not estimated within the age group.")

H2("3.6. Phân bố theo khu vực địa lý (Hình 3)")
nr("Tỉ lệ DALY chuẩn hoá theo tuổi do rối loạn tâm thần CAO NHẤT ở:")
nr("• Hoa Kỳ, Úc, New Zealand, Brazil")
nr("• Một số nơi ở Tây Âu: Greenland, Bồ Đào Nha, Hy Lạp, Ireland, Tây Ban Nha")
nr("• Vùng Tiểu Sahara: Uganda")
nr("• Bắc Phi và Trung Đông: Palestine, Lebanon, Iran")
nr("Tỉ lệ DALY chuẩn hoá THẤP NHẤT ở:")
nr("• Đông Nam Á: VIỆT NAM, Myanmar, Indonesia")
nr("• Đông Á: Đài Loan (Trung Quốc), Trung Quốc, Bắc Triều Tiên")
nr("• Châu Á - Thái Bình Dương thu nhập cao: Brunei")
nr("• Trung Á: Azerbaijan")
nr("LƯU Ý CÁCH ĐỌC: Bảng dưới đây là TỈ LỆ HIỆN MẮC chuẩn hoá theo tuổi (per 100.000) "
   "— tái tạo từ Bảng 2 PDF gốc trang 143-144. Đây là tỉ lệ HIỆN MẮC, KHÔNG phải tỉ lệ "
   "DALY. Tỉ lệ DALY toàn cầu của rối loạn tâm thần năm 2019 chỉ là 1.566,2/100.000 "
   "(Bảng 1 PDF). Em để bảng tỉ lệ hiện mắc theo vùng vì PDF gốc không tách bảng DALY "
   "theo 21 vùng — chỉ thể hiện qua Hình 3 bản đồ. Số liệu DALY toàn cầu xem mục 3.3.")
tbl(['Vùng', 'Tỉ lệ HIỆN MẮC 2019 chuẩn hoá theo tuổi (trên 100.000)'], [
    ['Toàn cầu', '12.262,0 (95% UI 11.382,9–13.213,3)'],
    ['Australasia (Úc + NZ)', '15.945,9 (12.875,7–17.139,0)'],
    ['Mỹ Latin nhiệt đới', '14.882,5 (12.105,7–16.946,5)'],
    ['Bắc Mỹ thu nhập cao', '15.909,4 (12.875,7–17.139,0)'],
    ['Đông Á', '10.566,7 (9.749,8–11.424,3)'],
    ['ĐÔNG NAM Á (chứa Việt Nam)', '10.520,4 (9.722,0–11.377,3)'],
    ['Châu Á-TBD thu nhập cao', '11.806,9 (10.916,6–12.775,4)'],
    ['Tiểu Sahara — Tây Phi', '11.000,7 (10.217,9–11.866,6)'],
    ['Trung Á', '10.179,5 (9.344,1–11.003,9)'],
], [4.5, 11.5])

insert_image('GBD2019_Figure3_world_map.png',
             'Hình 3 (trang 147 PDF gốc) — Bản đồ thế giới: tỉ lệ DALY chuẩn hoá theo tuổi '
             'trên 100.000 dân do rối loạn tâm thần, năm 2019. Càng đỏ = tỉ lệ càng cao; '
             'xanh dương = tỉ lệ thấp (Việt Nam thuộc nhóm thấp nhất)',
             width_cm=15)

H2("3.7. Rối loạn ăn uống — duy nhất có YLL")
nr("Năm 2019: rối loạn ăn uống là nguyên nhân của 318,3 ca tử vong toàn cầu (95% UI "
   "285,7–386,0) — chán ăn tâm thần (anorexia nervosa) chiếm phần lớn (268,7 ca; "
   "95% UI 242,5–326,9). Cuồng ăn tâm thần (bulimia nervosa) chiếm 49,6 ca (36,4–72,2).")
nr("CỰC THẤP nếu so với tử vong do các nguyên nhân khác — phản ánh hạn chế của khung "
   "GBD trong việc nắm bắt tử vong sớm của người có rối loạn tâm thần (xem mục 4.3 + "
   "phản biện).")

# ========== PHẦN 4 — THẢO LUẬN ==========
H1("PHẦN 4 — THẢO LUẬN")
page_marker('Trang 144-148, Lancet Psychiatry, Tập 9(2), 02/2022')

H2("4.1. Phát hiện chính được nhắc lại")
nr("• Rối loạn trầm cảm + lo âu là 2 nguyên nhân hàng đầu của gánh nặng SKTT toàn "
   "cầu — xếp hạng 13 (DALY) và 24 (YLD) trong 25 nguyên nhân hàng đầu, với tỉ lệ "
   "hiện mắc + trọng số khuyết tật cao hơn nhiều bệnh khác.")
nr("• Tâm thần phân liệt ảnh hưởng tỉ lệ NHỎ hơn dân số toàn cầu so với trầm cảm + "
   "lo âu, NHƯNG trọng số khuyết tật cho trạng thái loạn thần cấp (acute psychosis) "
   "là CAO NHẤT trong tất cả nghiên cứu GBD.")
nr("• Tỉ lệ rối loạn lo âu, trầm cảm, ăn uống lưỡng cực + ăn uống đặc biệt đáng quan "
   "tâm vì chúng KHÔNG CHỈ làm suy giảm sức khỏe trực tiếp, mà còn tăng nguy cơ các "
   "kết cục y tế khác như tự tử (xếp hạng nguyên nhân tử vong THỨ 18 toàn cầu trong "
   "GBD 2019).")

H2("4.2. Khác biệt theo giới + tuổi")
nr("• KHÔNG có khác biệt rõ rệt theo giới cho rối loạn lưỡng cực + tâm thần phân liệt")
nr("• NỮ cao hơn ở: rối loạn trầm cảm, lo âu, ăn uống")
nr("• NAM cao hơn ở: rối loạn phổ tự kỷ + ADHD")
nr("• Năm 2019: 80,6% gánh nặng SKTT xảy ra ở tuổi lao động (16-65) — đặt thách thức "
   "ĐÁNG KỂ cho nền kinh tế")
nr("• 23,2% trẻ em + vị thành niên có gánh nặng SKTT trên toàn cầu nằm ở Tiểu Sahara — "
   "nơi nguồn lực dành cho SKTT phát triển còn hạn chế")

H2("4.3. Vấn đề dữ liệu thiếu ở nước thu nhập thấp + trung bình")
nr("Tỉ lệ DALY do rối loạn tâm thần CAO ở nhiều nước thu nhập cao và THẤP NHẤT ở các "
   "vùng Tiểu Sahara + Châu Á — nhưng đây là nơi DỮ LIỆU DỊCH TỄ ÍT NHẤT → ƯỚC TÍNH "
   "có ĐỘ KHÔNG CHẮC CHẮN cao hơn. Tác giả khuyên: cần thêm khảo sát ở các nước này.")

H2("4.4. Tăng 55,1% tổng DALY 1990–2019")
nr("Tỉ lệ DALY chuẩn hoá theo tuổi GẦN NHƯ KHÔNG ĐỔI giữa 1990–2019 (từ 1581 lên "
   "1566/100.000), NHƯNG TỔNG SỐ DALY TĂNG 55,1% — chủ yếu do TĂNG DÂN SỐ + DÂN SỐ "
   "GIÀ HOÁ. Sự tăng này dự kiến TIẾP TỤC + nhấn mạnh nhu cầu hệ thống y tế (đặc "
   "biệt ở các nước thu nhập thấp + trung bình) cung cấp điều trị + chăm sóc cho "
   "dân số đang TĂNG NHANH.")

H2("4.5. Khoảng cách điều trị")
nr("Mặc dù có các gói can thiệp HIỆU QUẢ tồn tại, ở cấp toàn cầu vẫn có thiếu hụt "
   "ĐÁNG KỂ về tiếp cận dịch vụ và nguồn lực để mở rộng. Các rào cản đến chăm sóc "
   "bao gồm: thiếu nhu cầu cảm nhận (perceived need) + KỲ THỊ liên quan đến vấn đề "
   "SKTT. Ngay cả ở các nước thu nhập cao có tăng tiếp cận điều trị từ 1990, điều "
   "trị VẪN không đạt chuẩn tối thiểu hoặc không tới được người cần nhất.")

H2("4.6. Tác động COVID-19")
nr("Đại dịch COVID-19 nổi lên năm 2020 đã tạo môi trường có nhiều yếu tố gây xấu sức "
   "khỏe tâm thần. Tác động trực tiếp về tâm lý + tác động dài hạn về kinh tế-xã hội "
   "có thể TĂNG tỉ lệ rối loạn tâm thần phổ biến. Đã có nỗ lực ước tính tác động "
   "COVID-19 lên gánh nặng SKTT trong khuôn khổ GBD (sẽ được công bố riêng — Santomauro "
   "et al. 2021 trong Lancet).")

# ========== PHẦN 5 — REFERENCE ==========
H1("PHẦN 5 — DANH MỤC TÀI LIỆU THAM KHẢO TRONG BÀI GỐC")
nr("Bài gốc GBD 2019 Mental Disorders Collaborators (2022) có 37 tài liệu tham khảo. "
   "Em liệt kê 20 tài liệu đầu — phần em đã đối chiếu trực tiếp từ PDF.",
   italic=True, color=GRAY)
nr("(Giữ nguyên tiếng Anh — định dạng gốc của tạp chí Lancet Psychiatry)",
   italic=True, color=GRAY)
refs = [
    "Patel V, Saxena S, Lund C, et al. The Lancet Commission on global mental health and sustainable development. Lancet 2018; 392: 1553–98.",
    "Thornicroft G, Semrau M. Health system strengthening for mental health in low- and middle-income countries: introduction to the Emerald programme. BJPsych Open 2019; 5: e66.",
    "Wang H, Abbas KM, Abbasifard M, et al. Global age-sex-specific fertility, mortality, healthy life expectancy (HALE), and population estimates in 204 countries and territories, 1950–2019: a comprehensive demographic analysis for the Global Burden of Disease Study 2019. Lancet 2020; 396: 1160–203.",
    "Murray CJL, Aravkin AY, Zheng P, et al. Global burden of 87 risk factors in 204 countries and territories, 1990–2019: a systematic analysis for the Global Burden of Disease Study 2019. Lancet 2020; 396: 1223–49.",
    "Vos T, Lim SS, Abbafati C, et al. Global burden of 369 diseases and injuries in 204 countries and territories, 1990–2019: a systematic analysis for the Global Burden of Disease Study 2019. Lancet 2020; 396: 1204–22.",
    "Whiteford HA, Degenhardt L, Rehm J, et al. Global burden of disease attributable to mental and substance use disorders: findings from the Global Burden of Disease Study 2010. Lancet 2013; 382: 1575–86.",
    "Degenhardt L, Charlson F, Ferrari A, et al. The global burden of disease attributable to alcohol and drug use in 195 countries and territories, 1990-2016: a systematic analysis for the Global Burden of Disease Study 2016. Lancet Psychiatry 2018; 5: 987–1012.",
    "Stevens GA, Alkema L, Black RE, et al. Guidelines for Accurate and Transparent Health Estimates Reporting: the GATHER statement. Lancet 2016; 388: e19–23.",
    "Moffitt TE, Caspi A, Taylor A, et al. How common are common mental disorders? Evidence that lifetime prevalence rates are doubled by prospective versus retrospective ascertainment. Psychol Med 2010; 40: 899–909.",
    "Zheng P, Barber R, Sorensen RJD, Murray CJL, Aravkin AY. Trimmed constrained mixed effects models: formulations and algorithms. J Comput Graph Stat 2021; 30: 544–56.",
    "Flaxman AD, Vos T, Murray CJL. An integrative metaregression framework for descriptive epidemiology. Washington: University of Washington Press, 2015.",
    "Charlson F, van Ommeren M, Flaxman A, Cornett J, Whiteford H, Saxena S. New WHO prevalence estimates of mental disorders in conflict settings: a systematic review and meta-analysis. Lancet 2019; 394: 240–48.",
    "Gallup G. The Gallup poll: public opinion 2003. Lanham, MD: Rowman & Littlefield, 2004.",
    "Erskine HE, Ferrari AJ, Polanczyk GV, et al. The global burden of conduct disorder and attention-deficit/hyperactivity disorder in 2010. J Child Psychol Psychiatry 2014; 55: 328–36.",
    "Ferrari AJ, Saha S, McGrath JJ, et al. Health states for schizophrenia and bipolar disorder within the Global Burden of Disease 2010 Study. Popul Health Metr 2012; 10: 16.",
    "National Institute on Alcohol Abuse and Alcoholism. Introduction to the national epidemiologic survey on alcohol and related conditions 2016.",
    "Australian Bureau of Statistics. Mental Health and Wellbeing: Profile of Adults, Australia, 1997.",
    "Burstein R, Fleming T, Haagsma J, Salomon JA, Vos T, Scott JG. Estimating distributions of health state severity for the global burden of disease study. Popul Health Metr 2015; 13: 31.",
    "Baxter AJ, Brugha TS, Erskine HE, Scheurer RW, Vos T, Scott JG. The epidemiology and global burden of autism spectrum disorders. Psychol Med 2015; 45: 601–13.",
    "Baxter AJ, Scott KM, Vos T, Whiteford HA. Global prevalence of anxiety disorders: a systematic review and meta-regression. Psychol Med 2013; 43: 897–910.",
]
for i, r in enumerate(refs, 1):
    nr(f"{i}. {r}", size=11)
nr("(Còn 17 tài liệu tham khảo nữa — thầy có thể truy cập đầy đủ tại Lancet Psychiatry "
   "trang 149-150 PDF gốc)", italic=True, color=GRAY, size=10)

# ========== PHẦN 6 — PHẢN BIỆN ==========
H1("PHẦN 6 — PHẢN BIỆN CHI TIẾT")

H2("Điểm 1 — QUY MÔ CỰC LỚN + dữ liệu chuẩn quốc tế là điểm mạnh chính")
crit_para("GBD 2019 phân tích 12 rối loạn × 23 nhóm tuổi × 2 giới × 204 quốc gia × 30 "
          "năm — quy mô KHÔNG có nghiên cứu nào tương đương. Dùng phương pháp DisMod-MR "
          "2.1 — chuẩn vàng cho gộp dữ liệu dịch tễ. Có hệ thống tổng quan dữ liệu rõ "
          "ràng (Guidelines for Accurate and Transparent Health Estimates Reporting — "
          "GATHER) — đảm bảo minh bạch + có thể tái lập. Là CƠ SỞ DỮ LIỆU TOÀN CẦU "
          "CHUẨN — được trích trong hầu hết các bài về dịch tễ SKTT toàn cầu sau 2022 "
          "(bao gồm V-NAMHS 2022, Jenkins 2023 — đã có trong kho dự án).")

H2("Điểm 2 — Hạn chế CHÍNH: YLL CỰC THẤP cho rối loạn tâm thần")
crit_para("Tác giả TỰ thừa nhận trong Discussion: \"YLL ước tính cho rối loạn tâm thần "
          "là cực thấp và không phản ánh tử vong sớm thực tế ở người có rối loạn tâm "
          "thần\". Khuôn khổ GBD chỉ cho phép gán MỖI ca tử vong cho 1 nguyên nhân nền "
          "tảng theo ICD → tử vong do tự tử / tim mạch / hô hấp ở người có trầm cảm "
          "không được tính là YLL trầm cảm. Phân tích bổ sung GBD 2010 ước tính: nếu "
          "tính tự tử thuộc rối loạn tâm thần + lạm dụng chất, gánh nặng SKTT có thể "
          "TĂNG TỪ 7,4% lên 8,3% tổng DALY toàn cầu (xếp hạng từ thứ 5 lên thứ 3). Đây "
          "là vấn đề CỐT LÕI khi diễn giải số liệu GBD 2019.")

H2("Điểm 3 — Dữ liệu THIẾU ở các nước thu nhập thấp/trung bình + Việt Nam")
crit_para("VIỆT NAM được xếp vào nhóm có gánh nặng rối loạn tâm thần THẤP NHẤT thế "
          "giới theo Hình 3 PDF gốc (bản đồ Tỉ lệ DALY chuẩn hoá theo tuổi 2019). "
          "Tỉ lệ HIỆN MẮC chuẩn hoá theo tuổi của Đông Nam Á — vùng chứa Việt Nam — "
          "là 10.520,4/100.000 (cùng nhóm Myanmar, Indonesia, Trung Quốc). NHƯNG "
          "đây có thể KHÔNG phản ánh thực tế: do thiếu dữ liệu khảo sát chất lượng "
          "cao ở Việt Nam (V-NAMHS 2022 là khảo sát quốc gia ĐẦU TIÊN dùng công cụ "
          "DSM-5), DisMod-MR 2.1 phải dùng các đồng biến gián tiếp + ngoại suy từ "
          "vùng → ước tính có thể CHỆNH so với thực tế. So sánh: Hoa Kỳ + Úc dữ "
          "liệu phong phú → ước tính rất tin cậy. Cần thận trọng khi diễn giải \"Việt "
          "Nam có gánh nặng SKTT thấp\" — có thể chỉ là chưa có dữ liệu để xác nhận.")

H2("Điểm 4 — Định nghĩa case theo DSM-IV-TR / ICD-10 — đã LẠC HẬU")
crit_para("Bài dùng định nghĩa case theo DSM-IV-TR + ICD-10 vì đa số khảo sát toàn cầu "
          "dùng. Nhưng DSM-5 đã ban hành 2013, ICD-11 2019 — có nhiều thay đổi (vd: "
          "DSM-5 gộp Asperger vào rối loạn phổ tự kỷ; có thêm rối loạn nuốt chửng "
          "binge-eating). GBD 2019 KHÔNG nắm bắt các thay đổi này → có thể UNDERESTIMATE "
          "binge-eating disorder + các rối loạn ăn uống khác. Tác giả tự thừa nhận: "
          "\"work to account for the impact of changes to diagnostic classifications "
          "within our GBD estimates will be undertaken\".")

H2("Điểm 5 — Trọng số khuyết tật từ vài quốc gia chủ yếu thu nhập cao")
crit_para("Trọng số khuyết tật được tính từ khảo sát ở 9 quốc gia: Bangladesh, Indonesia, "
          "Peru, Tanzania, Hoa Kỳ, Hungary, Italia, Thuỵ Điển, Hà Lan + 1 khảo sát "
          "online tiếng Anh/Tây Ban Nha/Trung. Áp dụng trọng số này CHO TẤT CẢ 204 "
          "quốc gia → giả định văn hoá tương đồng về cách CẢM NHẬN khuyết tật. Tuy "
          "nhiên hiểu biết về SKTT khác nhau theo văn hoá (cùng triệu chứng nhưng "
          "trọng số khuyết tật cảm nhận có thể khác). Tác giả tự thừa nhận: \"Imposing "
          "severity distributions from high-income locations to all locations is "
          "likely to have underestimated burden in countries with little or no access "
          "to treatment\".")

# ========== PHẦN 7 — GAP ==========
H1("PHẦN 7 — PHÁT HIỆN GAP NGHIÊN CỨU + CHO VIỆT NAM")

H2("Gap 1 — Việt Nam thiếu dữ liệu dịch tễ SKTT chất lượng cao trong giai đoạn dài")
nr("V-NAMHS 2022 là khảo sát quốc gia ĐẦU TIÊN của Việt Nam dùng công cụ chẩn đoán "
   "DSM-5 chuẩn (DISC-5). GBD 2019 khi ước tính cho Việt Nam đã PHẢI dựa vào dữ liệu "
   "khu vực + đồng biến gián tiếp. Cần thực hiện khảo sát quốc gia LẶP LẠI ít nhất "
   "5 năm/lần để có dữ liệu thời gian thực. Các tỉnh miền núi + DTTS cần khảo sát "
   "RIÊNG vì khả năng có gánh nặng KHÁC vùng đồng bằng.")

H2("Gap 2 — Việt Nam chưa tham gia GBD Collaborator Network chính thức")
nr("Trong danh sách ~150 tác giả của GBD 2019, có cộng tác viên từ Trung Quốc, Ấn Độ, "
   "Brazil, Mexico, Tây Ban Nha — KHÔNG có Việt Nam. Việc tham gia GBD Collaborator "
   "Network sẽ giúp Việt Nam: (a) cung cấp dữ liệu vào hệ thống toàn cầu; (b) học "
   "phương pháp DisMod-MR 2.1; (c) có ước tính chính xác hơn cho VN trong các cập "
   "nhật GBD tương lai (GBD 2024, GBD 2027).")

H2("Gap 3 — Cần phương pháp tính TỬ VONG do rối loạn tâm thần phù hợp hơn")
nr("YLL gần như bằng 0 cho hầu hết rối loạn tâm thần là HẠN CHẾ NGHIÊM TRỌNG. Cần "
   "phương pháp counterfactual (so sánh với người không có rối loạn) hoặc phương "
   "pháp \"comparative risk assessment\" để phân bổ phần tử vong sớm có thể quy cho "
   "rối loạn tâm thần. Đã có nỗ lực (Walker et al. 2015 JAMA Psychiatry; Plana-Ripoll "
   "et al. 2019 Lancet) — cần tích hợp vào GBD chính thức.")

H2("Gap 4 — Cập nhật theo DSM-5 + ICD-11")
nr("Cần cập nhật phân loại rối loạn theo DSM-5 + ICD-11 — đặc biệt cho binge-eating "
   "disorder (BED) + các \"other specified feeding/eating disorders\" (OSFED) hiện "
   "chưa được tách riêng. Cũng cần cập nhật cho các rối loạn nhân cách (personality "
   "disorders) — hiện gộp trong nhóm \"other mental disorders\".")

H2("Gap 5 — Cần kết nối GBD với các nỗ lực CAN THIỆP có hiệu quả")
nr("GBD chỉ ĐO gánh nặng — không đo HIỆU QUẢ can thiệp. Cần kết nối GBD với các phân "
   "tích chi phí-hiệu quả (cost-effectiveness) như WHO-CHOICE để đề xuất GÓI can thiệp "
   "ưu tiên cho từng quốc gia + từng nhóm tuổi. Việt Nam đặc biệt cần các phân tích "
   "này để xây dựng chiến lược quốc gia về SKTT.")

H2("Gap 6 — Cần tích hợp tác động COVID-19")
nr("Bài này dùng dữ liệu đến 2019 — không tính tác động COVID-19 (2020+). Đã có bài "
   "tiếp theo (Santomauro et al. 2021 trong Lancet) ước tính tăng 25% trầm cảm + lo "
   "âu toàn cầu năm 2020. Cần kết hợp 2 ước tính + dõi theo các năm sau (2021-2024) "
   "để đánh giá phục hồi (hoặc không).")

H2("Áp dụng cho VIỆT NAM — đề xuất cụ thể")
nr("(1) Dùng số liệu GBD 2019 làm THAM CHIẾU TOÀN CẦU + đối chiếu với V-NAMHS 2022 "
   "ở Việt Nam — phát hiện các khoảng trống đo lường.")
nr("(2) Đề xuất Bộ Y tế Việt Nam cử đại diện tham gia GBD Collaborator Network — "
   "kênh: Institute for Health Metrics and Evaluation, Đại học Washington, Hoa Kỳ.")
nr("(3) Xây dựng KHẢO SÁT QUỐC GIA SKTT VIỆT NAM định kỳ 5 năm/lần — bao quát: "
   "(a) trẻ em + vị thành niên (V-NAMHS 2022 đã làm); (b) người lớn (chưa có khảo "
   "sát quốc gia gần đây); (c) người cao tuổi.")
nr("(4) Trong giáo dục y khoa Việt Nam, dùng GBD 2019 làm tài liệu chuẩn dạy về dịch "
   "tễ SKTT toàn cầu — kết hợp với dữ liệu Việt Nam.")
nr("(5) Các tổ chức xã hội + truyền thông Việt Nam có thể trích GBD 2019 (cùng các "
   "cập nhật) để vận động cho ưu tiên SKTT trong chính sách y tế quốc gia.")

d.save(OUT)
print('Wrote:', OUT)
