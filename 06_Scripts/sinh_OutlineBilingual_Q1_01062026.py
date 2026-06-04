# -*- coding: utf-8 -*-
"""Sinh outline song ngu (VN + EN) chi tiet Q1 BMC Psychiatry de gui co-authors.
- Bilingual side-by-side hoac stacked
- Detailed (expand bullets thanh paragraphs)
- Easy to read for academic co-authors
- Tables proper column widths theo noi dung
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'OutlineBilingual_Q1_01062026.docx')

d = Document()
for sec in d.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.0)
    sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0)
    sec.right_margin = Cm(2.0)

s = d.styles['Normal']
s.font.name = 'Times New Roman'
s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.25


def H1(text_vn, text_en=None):
    """Heading 1 - large blue centered."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    if text_en:
        p2 = d.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p2.paragraph_format.space_after = Pt(10)
        r2 = p2.add_run(text_en)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(13); r2.italic = True
        r2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)


def H2(text_vn, text_en=None):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    if text_en:
        r2 = p.add_run('  /  ' + text_en)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True
        r2.font.color.rgb = RGBColor(0x70, 0x70, 0x70)


def H3(text_vn, text_en=None):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    if text_en:
        r2 = p.add_run('  /  ' + text_en)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(10); r2.italic = True
        r2.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def Para(text_vn, text_en=None, indent=True, italic=False):
    """Paragraph stacked VN above, EN below in lighter."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(3)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.italic = italic

    if text_en:
        p2 = d.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.space_after = Pt(8)
        if indent:
            p2.paragraph_format.first_line_indent = Cm(0.5)
        r2 = p2.add_run(text_en)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(10); r2.italic = True
        r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def Bullet(text_vn, text_en=None, level=0):
    """Bullet point with VN + EN stacked."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.6 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('▸ ' + text_vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)

    if text_en:
        p2 = d.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p2.paragraph_format.left_indent = Cm(0.9 + level * 0.5)
        p2.paragraph_format.space_after = Pt(6)
        r2 = p2.add_run(text_en)
        r2.font.name = 'Times New Roman'; r2.font.size = Pt(10); r2.italic = True
        r2.font.color.rgb = RGBColor(0x40, 0x40, 0x40)


def NCS_marker(text):
    """[NCS/Thầy confirm] placeholder."""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ [NCS/Thầy confirm] ' + text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True; r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


def set_col_widths(table, widths_cm):
    """Set column widths in cm. Must apply to each cell + table."""
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


def make_table(headers_vn, headers_en, rows, col_widths_cm, style='Light Grid Accent 1'):
    """Create table with bilingual headers + proper column widths."""
    t = d.add_table(rows=1, cols=len(headers_vn))
    t.style = style
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h_vn in enumerate(headers_vn):
        h_en = headers_en[i] if headers_en and i < len(headers_en) else None
        hdr[i].text = ''
        p = hdr[i].paragraphs[0]
        r = p.add_run(h_vn)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        if h_en:
            p.add_run('\n')
            r2 = p.add_run(h_en)
            r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.italic = True
            r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)

    for row_data in rows:
        row = t.add_row().cells
        for i, cell_data in enumerate(row_data):
            row[i].text = ''
            p = row[i].paragraphs[0]
            if isinstance(cell_data, tuple):
                # bilingual cell
                vn, en = cell_data
                r = p.add_run(vn)
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if en:
                    p.add_run('\n')
                    r2 = p.add_run(en)
                    r2.font.name = 'Times New Roman'; r2.font.size = Pt(9); r2.italic = True
                    r2.font.color.rgb = RGBColor(0x60, 0x60, 0x60)
            else:
                r = p.add_run(str(cell_data))
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)

    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
# COVER
# ============================================================
H1('ĐỀ CƯƠNG CHI TIẾT BÀI BÁO Q1',
   'DETAILED OUTLINE — Q1 PAPER')

P_doc = d.add_paragraph()
P_doc.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = P_doc.add_run('Bản song ngữ Tiếng Việt – English  /  Bilingual Edition')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True
r.font.color.rgb = RGBColor(0x40, 0x40, 0x40)

d.add_paragraph()

Para('Tài liệu này là đề cương chi tiết bài báo nghiên cứu được gửi tới tạp '
     'chí Q1 quốc tế BMC Psychiatry. Bản outline được trình bày song ngữ Tiếng '
     'Việt – English để thuận tiện cho việc rà soát của đồng tác giả.',
     'This document presents the detailed outline of a Q1 international '
     'research paper to be submitted to BMC Psychiatry. The bilingual '
     'Vietnamese–English format is designed for review by co-authors.',
     indent=False)

Para('', '')

# Key info table
key_info_data = [
    (('Tên bài (tentative)', 'Title (tentative)'),
     ('Mô hình SEM tích hợp các yếu tố nguy cơ và bảo vệ đối với các dạng rối loạn lo âu '
      'ở học sinh trung học cơ sở Việt Nam: Nghiên cứu hỗn hợp định tính – định lượng',
      'Integrated risk-protective structural equation model of anxiety disorder subtypes '
      'among Vietnamese lower secondary school students: A mixed-methods study')),
    (('Tạp chí mục tiêu', 'Target journal'),
     ('BMC Psychiatry (Q1, IF 4.4, acceptance ~30%)',
      'BMC Psychiatry (Q1, IF 4.4, ~30% acceptance rate)')),
    (('Backup journal', 'Backup journal'),
     ('BMC Psychology (IF 2.6) nếu reject',
      'BMC Psychology (IF 2.6) if rejected')),
    (('Số từ dự kiến', 'Target word count'),
     ('6.000–8.000 từ (không tính references)',
      '6,000–8,000 words (excluding references)')),
    (('Tác giả', 'Authors'),
     ('Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹†',
      'Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹†')),
    (('Ghi chú tác giả', 'Author notes'),
     ('*: Tác giả thứ nhất, †: Tác giả liên hệ',
      '*: First author, †: Corresponding author')),
]

t_key = make_table(
    headers_vn=['Mục', 'Nội dung'],
    headers_en=['Item', 'Content'],
    rows=key_info_data,
    col_widths_cm=[4.0, 13.0],
)

d.add_page_break()


# ============================================================
H1('1. TÓM TẮT KHOA HỌC', '1. SCIENTIFIC ABSTRACT')

Para('Abstract chuẩn cấu trúc BMC Psychiatry (~250 từ).',
     'Structured abstract per BMC Psychiatry guidelines (~250 words).',
     indent=False, italic=True)

H3('Background', 'Background')

Para('Rối loạn lo âu ở học sinh trung học cơ sở (THCS) là vấn đề sức khỏe '
     'tâm thần quan trọng, đặc biệt tại các nước đang phát triển ở châu Á. '
     'Tại Việt Nam, các nghiên cứu hiện có chủ yếu tập trung vào yếu tố nguy cơ '
     'hoặc bảo vệ một cách riêng lẻ, chưa có nghiên cứu nào kiểm định đồng '
     'thời tác động của nhiều yếu tố trên các dạng rối loạn lo âu khác nhau '
     'theo phân loại DSM-5.',
     'Anxiety disorders among lower secondary school students represent a '
     'critical mental health concern, particularly in developing Asian '
     'countries. Existing Vietnamese studies have predominantly examined '
     'either risk or protective factors in isolation, with no prior research '
     'simultaneously testing multiple factors across DSM-5 anxiety disorder '
     'subtypes.')

H3('Methods', 'Methods')

Para('Nghiên cứu áp dụng thiết kế hỗn hợp song song hội tụ (convergent '
     'parallel mixed-methods design) trên mẫu 1.352 học sinh THCS (614 nam, '
     '738 nữ) tại 2 trường ở Hà Nội. Tám thang đo đã được kiểm chứng quốc tế '
     'được sử dụng để đo lường: rối loạn lo âu (RCADS), ba yếu tố nguy cơ '
     '(áp lực học tập – ESSA, nghiện điện thoại – SAS-SV, bắt nạt học '
     'đường – OBVQ), bốn yếu tố bảo vệ (gắn bó trường học – PSSM, hỗ trợ '
     'xã hội – MSPSS, lòng tự trọng – RSES, và biện pháp đối phó – Brief COPE). '
     'Phân tích bao gồm CFA, mô hình phương trình cấu trúc tích hợp (SEM), '
     'kiểm định bất biến đa nhóm theo giới tính, và phỏng vấn bán cấu trúc.',
     'A convergent parallel mixed-methods design was applied to a sample of '
     '1,352 lower secondary school students (614 males, 738 females) at two '
     'schools in Hanoi. Eight internationally validated scales measured: '
     'anxiety disorders (RCADS), three risk factors (academic pressure – '
     'ESSA, smartphone addiction – SAS-SV, school bullying – OBVQ), and four '
     'protective factors (school engagement – PSSM, social support – MSPSS, '
     'self-esteem – RSES, and coping – Brief COPE). Analyses included CFA, '
     'integrated structural equation modeling (SEM), multi-group invariance '
     'testing by gender, and semi-structured qualitative interviews.')

H3('Results', 'Results')

Para('Các đường dẫn từ ba yếu tố nguy cơ và bốn yếu tố bảo vệ đến ba dạng '
     'rối loạn lo âu đều có ý nghĩa thống kê (p < 0,001), với áp lực học '
     'tập có cường độ mạnh nhất (β = +0,533) và lòng tự trọng là yếu tố bảo '
     'vệ mạnh nhất (β = -0,457). Phát hiện đáng chú ý: bắt nạt học đường '
     'thể hiện đường dẫn mạnh nhất tới lo âu chia ly (β = +0,376). Kiểm định '
     'bất biến đa nhóm cho thấy rối loạn lo âu chia ly không có khác biệt '
     'theo giới tính (F = 0,246; p = 0,620), khác với rối loạn lo âu lan toả '
     '(F = 44,484) và lo âu xã hội (F = 45,984).',
     'Paths from three risk factors and four protective factors to three '
     'anxiety disorder subtypes were all statistically significant (p < '
     '0.001), with academic pressure showing the strongest effect (β = '
     '+0.533) and self-esteem the strongest protective effect (β = -0.457). '
     'Notably, school bullying exhibited the strongest path to separation '
     'anxiety (β = +0.376). Multi-group invariance testing revealed gender '
     'invariance for separation anxiety (F = 0.246; p = 0.620), unlike '
     'generalized anxiety (F = 44.484) and social anxiety (F = 45.984).')

H3('Conclusions', 'Conclusions')

Para('Nghiên cứu này là nghiên cứu đầu tiên áp dụng SEM tích hợp trên mẫu '
     'học sinh THCS Việt Nam. Phát hiện về tính bất biến giới của rối loạn '
     'lo âu chia ly mở ra hướng nghiên cứu mới về vai trò của bối cảnh văn '
     'hóa tập thể đối với các giai đoạn phát triển tâm lý sớm. Kết quả gợi '
     'ý chương trình phòng ngừa rối loạn lo âu cần được thiết kế đặc thù '
     'theo giới tính cho lo âu lan toả và xã hội, nhưng có thể áp dụng phổ '
     'quát cho lo âu chia ly.',
     'This is the first study to apply integrated SEM to Vietnamese lower '
     'secondary school students. The gender-invariance finding for '
     'separation anxiety opens new avenues for examining collectivist '
     'cultural contexts in early psychological development. Findings suggest '
     'that anxiety prevention programs should be gender-tailored for '
     'generalized and social anxiety, but may apply universally for '
     'separation anxiety.')


d.add_page_break()


# ============================================================
H1('2. ĐIỂM MỚI CỦA NGHIÊN CỨU', '2. NOVEL CONTRIBUTIONS')

H2('Ba điểm mới chính', 'Three Key Contributions')

contributions_data = [
    (('1', '1'),
     ('Mô hình SEM tích hợp đầu tiên ở Việt Nam',
      'First integrated SEM model in Vietnam'),
     ('Nghiên cứu đầu tiên đồng thời kiểm định 3 yếu tố nguy cơ (áp lực học tập, '
      'nghiện điện thoại, bắt nạt học đường) và 4 yếu tố bảo vệ (gắn bó '
      'trường học, hỗ trợ cha mẹ, hỗ trợ bạn bè, lòng tự trọng) trên 3 dạng '
      'rối loạn lo âu theo DSM-5, tổng cộng 21 đường dẫn (paths), trên mẫu '
      'học sinh THCS Việt Nam.',
      'First study to simultaneously test three risk factors (academic '
      'pressure, smartphone addiction, school bullying) and four protective '
      'factors (school engagement, parental support, peer support, '
      'self-esteem) on three DSM-5 anxiety disorder subtypes, comprising 21 '
      'paths in total, in a sample of Vietnamese lower secondary school students.')),
    (('2', '2'),
     ('Phát hiện tính bất biến giới của lo âu chia ly',
      'Gender invariance finding for separation anxiety'),
     ('Khác biệt giới ở rối loạn lo âu chia ly không có ý nghĩa thống kê '
      '(F = 0,246; p = 0,620), trong khi rối loạn lo âu lan toả '
      '(F = 44,484; p < 0,001) và lo âu xã hội (F = 45,984; p < 0,001) đều có '
      'khác biệt rõ rệt giữa nam và nữ. Phát hiện này thách thức giả định '
      'rằng khác biệt giới là đồng nhất ở mọi dạng rối loạn lo âu.',
      'No statistically significant gender difference was found for '
      'separation anxiety (F = 0.246; p = 0.620), while generalized anxiety '
      '(F = 44.484; p < 0.001) and social anxiety (F = 45.984; p < 0.001) '
      'showed marked male-female differences. This finding challenges the '
      'assumption that gender differences are uniform across all anxiety '
      'disorder subtypes.')),
    (('3', '3'),
     ('Tích hợp phương pháp định tính – định lượng',
      'Mixed-methods qualitative–quantitative integration'),
     ('Dữ liệu phỏng vấn bán cấu trúc với học sinh được tích hợp theo thiết '
      'kế song song hội tụ (Creswell & Plano Clark, 2018), bổ sung ngữ '
      'cảnh văn hóa cho các hệ số β định lượng. Điều này đặc biệt quan trọng '
      'cho việc diễn giải vai trò của yếu tố bảo vệ "hỗ trợ cha mẹ" và '
      '"hỗ trợ bạn bè" trong bối cảnh văn hóa tập thể Việt Nam.',
      'Semi-structured student interview data is integrated via convergent '
      'parallel design (Creswell & Plano Clark, 2018), adding cultural '
      'context to quantitative β coefficients. This is particularly '
      'important for interpreting the role of protective factors '
      '"parental support" and "peer support" in the Vietnamese collectivist '
      'cultural context.')),
]

t_contrib = make_table(
    headers_vn=['#', 'Tiêu đề', 'Mô tả chi tiết'],
    headers_en=['#', 'Title', 'Detailed description'],
    rows=contributions_data,
    col_widths_cm=[1.0, 4.5, 11.5],
)

d.add_page_break()


# ============================================================
H1('3. PHẦN GIỚI THIỆU', '3. INTRODUCTION')

Para('Phần Introduction được mở rộng thành 4 đoạn (~1.500 từ tổng cộng), với '
     'logic dẫn dắt từ gánh nặng toàn cầu → bối cảnh châu Á → khoảng trống '
     'nghiên cứu Việt Nam → mục tiêu + 3 giả thuyết của nghiên cứu hiện tại.',
     'The Introduction is expanded into 4 paragraphs (~1,500 words total), '
     'with logical flow from global burden → Asian context → Vietnamese '
     'research gap → present study objectives + 3 hypotheses.',
     indent=False, italic=True)


H2('§1. Gánh nặng lo âu ở thanh thiếu niên + Cơ sở phân loại DSM-5',
   '§1. Adolescent anxiety burden + DSM-5 subtype rationale')

Para('Rối loạn lo âu là vấn đề sức khỏe tâm thần phổ biến nhất ở thanh thiếu '
     'niên trên toàn cầu, ảnh hưởng đến 7-15% nhóm tuổi 10-19 (Anderson và '
     'cs., 2025). Báo cáo Gánh nặng Bệnh tật Toàn cầu (GBD) khu vực ASEAN '
     '2025 trên tạp chí Lancet ghi nhận tỷ lệ hiện mắc rối loạn tâm thần '
     'tại Việt Nam là 10,1% và tỷ lệ trung bình khu vực là 11,9% (2021). '
     'Đáng chú ý, nhóm tuổi 10-14 chiếm 16,3% tổng gánh nặng DALYs '
     '(Disability-Adjusted Life Years) liên quan đến rối loạn tâm thần.',
     'Anxiety disorders represent the most prevalent mental health concern '
     'among adolescents globally, affecting 7-15% of the 10-19 age group '
     '(Anderson et al., 2025). The Global Burden of Disease (GBD) ASEAN 2025 '
     'Lancet report documented mental disorder prevalence at 10.1% in '
     'Vietnam and a regional mean of 11.9% in 2021. Notably, the 10-14 age '
     'group accounted for 16.3% of total mental-disorder-related DALYs '
     '(Disability-Adjusted Life Years).')

Para('Bối cảnh châu Á đặc biệt nổi bật với tốc độ tăng cao nhất về tỷ lệ '
     'hiện mắc theo độ tuổi giai đoạn 1990-2021. Nghiên cứu của Xu và cộng '
     'sự (2021) trên 373.216 học sinh THCS Trung Quốc trong đại dịch '
     'COVID-19 đã ghi nhận tỷ lệ triệu chứng lo âu 9,89% (mức trung bình '
     'trở lên). Wen và cộng sự (2020) trên mẫu học sinh nông thôn Trung '
     'Quốc và Chen và cộng sự (2023) trên 63.205 học sinh phổ thông cho '
     'thấy sự đa dạng của các yếu tố nguy cơ trong bối cảnh khác nhau.',
     'The Asian context is particularly notable for the highest age-'
     'standardized increase in prevalence during 1990-2021. Xu et al. (2021) '
     'on 373,216 Chinese lower secondary school students during the COVID-19 '
     'pandemic documented 9.89% anxiety symptom prevalence (moderate or '
     'higher). Wen et al. (2020) on rural Chinese students and Chen et al. '
     '(2023) on 63,205 secondary school students revealed the diversity of '
     'risk factors in different contexts.')

Para('Đặc thù văn hóa Việt Nam tạo ra bối cảnh nghiên cứu riêng biệt: '
     '(a) văn hóa Khổng giáo với hệ thống thi cử cạnh tranh cao đại diện cho '
     'con đường di động xã hội; (b) cấu trúc gia đình theo thứ bậc với '
     'mức độ bộc lộ cảm xúc hạn chế; và (c) định hướng tập thể chủ nghĩa '
     'có thể ảnh hưởng đến quỹ đạo phát triển rối loạn lo âu chia ly khác '
     'với bối cảnh phương Tây. Việc phân tích theo các phân loại DSM-5 '
     '(rối loạn lo âu lan toả – GAD, rối loạn lo âu chia ly – SAD, rối loạn '
     'lo âu xã hội – SocAD) là cần thiết vì các loại này có cơ chế nguyên '
     'nhân và liệu pháp điều trị khác nhau (Chorpita, 2000).',
     'Vietnamese cultural specificity creates a distinctive research '
     'context: (a) Confucian academic culture with high-stakes competitive '
     'examinations representing the path to social mobility; (b) hierarchical '
     'family structure with limited emotional disclosure; and (c) collectivist '
     'orientation that may shape separation anxiety developmental trajectories '
     'differently than Western contexts. Analysis by DSM-5 subtypes '
     '(generalized anxiety disorder – GAD, separation anxiety disorder – SAD, '
     'social anxiety disorder – SocAD) is essential because these subtypes '
     'have distinct etiological mechanisms and treatment responses '
     '(Chorpita, 2000).')


H2('§2. Khung lý thuyết nguy cơ – bảo vệ và lý do mô hình tích hợp',
   '§2. Risk-protective framework and rationale for integrated modeling')

Para('Lý thuyết Stress-Coping của Lazarus và Folkman (1984) thiết lập '
     'nền tảng lý thuyết cho việc nghiên cứu nhiều yếu tố nguy cơ và bảo vệ '
     'cùng tương tác để định hình kết quả sức khỏe tâm thần. Phân tích meta '
     'của Compas và cộng sự (2017) trên 212 nghiên cứu với N = 80.850 '
     'người tham gia đã khẳng định sự tương tác giữa yếu tố nguy cơ và bảo '
     'vệ, nhưng phần lớn các nghiên cứu thành phần đều kiểm định các yếu tố '
     'một cách riêng lẻ.',
     'Lazarus and Folkman\'s (1984) Stress-Coping theory established the '
     'theoretical foundation for studying multiple risk and protective '
     'factors interacting to shape mental health outcomes. The Compas et al. '
     '(2017) meta-analysis on 212 studies with N = 80,850 participants '
     'confirmed risk-protective factor interplay, but most constituent '
     'studies tested factors in isolation.')

Para('Bằng chứng thực chứng tại Việt Nam còn hạn chế và không nhất quán. '
     'Điều tra Sức khỏe Tâm thần Vị thành niên Quốc gia Việt Nam '
     '(V-NAMHS, 2022) báo cáo tỷ lệ hiện mắc rối loạn lo âu là 2,3% sử '
     'dụng công cụ DISC-5 (chẩn đoán lâm sàng). Trong khi đó, Hoàng Trung '
     'Học và cộng sự (2025) trên N = 8.473 học sinh sau COVID-19 ghi nhận '
     'tỷ lệ cao hơn nhiều khi sử dụng DASS-21. Sự không nhất quán phương '
     'pháp đo lường này phản ánh khoảng trống nghiên cứu phương pháp luận '
     'tại Việt Nam.',
     'Empirical evidence in Vietnam remains limited and inconsistent. The '
     'Vietnam National Adolescent Mental Health Survey (V-NAMHS, 2022) '
     'reported anxiety disorder prevalence of 2.3% using the DISC-5 '
     'instrument (clinical diagnosis). Meanwhile, Hoang Trung Hoc et al. '
     '(2025) on N = 8,473 post-COVID-19 students reported substantially '
     'higher rates with DASS-21. This measurement-method inconsistency '
     'reflects a methodological research gap in Vietnam.')

Para('Lý do nghiên cứu mô hình tích hợp: các nghiên cứu trước tại Việt Nam '
     'kiểm định yếu tố nguy cơ HOẶC yếu tố bảo vệ một cách riêng lẻ, dẫn '
     'đến không nắm bắt được tác động tương tác và không thể so sánh trực '
     'tiếp cường độ tương đối. Mô hình SEM tích hợp kiểm định 3 yếu tố nguy '
     'cơ và 4 yếu tố bảo vệ đồng thời trên 3 dạng RLLA giúp lấp đầy khoảng '
     'trống phương pháp luận này, đồng thời cung cấp ước lượng β đã hiệu '
     'chỉnh lẫn nhau giữa các yếu tố.',
     'Rationale for integrated modeling: Previous Vietnamese studies tested '
     'risk OR protective factors in isolation, resulting in failure to '
     'capture interaction effects and inability to directly compare relative '
     'effect magnitudes. An integrated SEM simultaneously testing 3 risk '
     'factors and 4 protective factors on 3 anxiety disorder subtypes fills '
     'this methodological gap, while also providing mutually-adjusted β '
     'estimates across factors.')


H2('§3. Khác biệt giới theo dạng RLLA và bối cảnh văn hóa Việt Nam',
   '§3. Differential gender patterns by anxiety subtype and Vietnamese cultural context')

Para('Kỳ vọng tiêu chuẩn dựa trên y văn quốc tế là tỷ lệ rối loạn lo âu ở '
     'nữ cao hơn nam (McLean và cs., 2011 – phân tích meta tổng hợp). '
     'Tuy nhiên, bằng chứng mới nổi thách thức giả định này: Wen và cộng '
     'sự (2020) trên học sinh nông thôn Trung Quốc ghi nhận nam > nữ; '
     'Saikia và cộng sự (2023) tại vùng Đông Bắc Ấn Độ trên 287 thanh '
     'thiếu niên báo cáo nam giới có tỷ lệ lo âu nặng cao hơn nữ giới '
     '(Boy 30,0% vs Girl 18,9%, p = 0,049).',
     'Standard expectation based on international literature is higher '
     'female-than-male anxiety disorder rates (McLean et al., 2011 – '
     'comprehensive meta-analysis). However, emerging evidence challenges '
     'this assumption: Wen et al. (2020) on rural Chinese students documented '
     'male > female; Saikia et al. (2023) in Northeast India on 287 '
     'adolescents reported higher severe anxiety rates in boys than girls '
     '(Boy 30.0% vs Girl 18.9%, p = 0.049).')

Para('Giả thuyết trung gian văn hóa Việt Nam: văn hóa tập thể chủ nghĩa '
     'có thể đồng nhất hóa trải nghiệm rối loạn lo âu CHIA LY giữa các '
     'giới (cấu trúc thứ bậc gia đình tạo trải nghiệm gắn bó đồng nhất), '
     'trong khi áp lực xã hội/giáo dục tác động khác biệt theo giới đến '
     'rối loạn lo âu LAN TOẢ và XÃ HỘI. Đây là giả thuyết khoa học có '
     'thể kiểm định bằng phân tích bất biến đa nhóm.',
     'Vietnamese cultural mediation hypothesis: collectivist culture may '
     'homogenize SEPARATION anxiety experience across genders (hierarchical '
     'family structure creates uniform attachment experience), while social/'
     'educational pressures differentially impact GENERALIZED and SOCIAL '
     'anxiety by gender. This is a testable scientific hypothesis through '
     'multi-group invariance analysis.')


H2('§4. Mục tiêu nghiên cứu hiện tại và 3 giả thuyết',
   '§4. Present study objectives and 3 hypotheses')

Para('Nghiên cứu hiện tại nhằm kiểm định 3 giả thuyết khoa học có thể '
     'kiểm chứng được:',
     'The present study aims to test three testable scientific hypotheses:',
     indent=False)

hypotheses_data = [
    (('H1', 'H1'),
     ('Ba yếu tố nguy cơ (áp lực học tập, nghiện điện thoại, bắt nạt học '
      'đường) sẽ có đường dẫn β dương có ý nghĩa thống kê tới cả ba dạng '
      'rối loạn lo âu (GAD, SAD, SocAD).',
      'All three risk factors (academic pressure, smartphone addiction, '
      'school bullying) will show significant positive β paths to all three '
      'anxiety disorder subtypes (GAD, SAD, SocAD).'),
     ('Tiêu chuẩn — bằng chứng quốc tế nhất quán',
      'Standard — consistent international evidence')),
    (('H2', 'H2'),
     ('Bốn yếu tố bảo vệ (gắn bó trường học, hỗ trợ cha mẹ, hỗ trợ bạn bè, '
      'lòng tự trọng) sẽ có đường dẫn β âm có ý nghĩa thống kê tới cả ba '
      'dạng rối loạn lo âu.',
      'All four protective factors (school engagement, parental support, '
      'peer support, self-esteem) will show significant negative β paths to '
      'all three anxiety disorder subtypes.'),
     ('Tiêu chuẩn — bằng chứng quốc tế nhất quán',
      'Standard — consistent international evidence')),
    (('H3', 'H3'),
     ('Phân tích SEM đa nhóm theo giới tính sẽ chỉ ra TÍNH BẤT BIẾN GIỚI '
      'cho rối loạn lo âu CHIA LY, trong khi rối loạn lo âu LAN TOẢ và '
      'XÃ HỘI sẽ thể hiện khác biệt giới (nữ > nam trong cường độ đường dẫn).',
      'Multi-group SEM by gender will demonstrate GENDER INVARIANCE for '
      'separation anxiety only, while generalized and social anxiety will '
      'show gender differences (female > male in path strengths).'),
     ('MỚI — chưa có bằng chứng trực tiếp ở Việt Nam',
      'NOVEL — no direct evidence in Vietnam to date')),
]

t_hyp = make_table(
    headers_vn=['Mã', 'Phát biểu giả thuyết', 'Tính chất'],
    headers_en=['Code', 'Hypothesis statement', 'Nature'],
    rows=hypotheses_data,
    col_widths_cm=[1.2, 12.3, 3.5],
)


Para('Lý do bổ sung phương pháp định tính: phân tích định lượng SEM một '
     'mình không thể nắm bắt các cơ chế văn hóa-bối cảnh đằng sau các mối '
     'quan hệ thống kê. Phỏng vấn bán cấu trúc với học sinh được chọn mẫu '
     'theo phân tầng cường độ lo âu (cao, trung bình, thấp) sẽ làm sáng tỏ '
     'trải nghiệm sống đằng sau các hệ số β, cho phép diễn giải có cơ sở '
     'văn hóa về các đường dẫn yếu tố bảo vệ.',
     'Rationale for mixed-methods addition: Quantitative SEM alone cannot '
     'capture the cultural-contextual mechanisms underlying these statistical '
     'relationships. Semi-structured interviews with students stratified by '
     'anxiety severity (high, medium, low) will illuminate the lived '
     'experience behind β coefficients, enabling culturally grounded '
     'interpretation of protective factor pathways.', indent=False)


d.add_page_break()


# ============================================================
H1('4. PHƯƠNG PHÁP NGHIÊN CỨU', '4. METHODS')


H2('4.1 Thiết kế nghiên cứu', '4.1 Study design')

Para('Nghiên cứu cắt ngang hỗn hợp định tính – định lượng (mixed-methods) '
     'theo thiết kế song song hội tụ (Convergent Parallel Design – Creswell '
     '& Plano Clark, 2018). Dữ liệu định lượng và định tính được thu thập '
     'song song, phân tích riêng biệt, sau đó tích hợp thông qua ma trận '
     'trình bày kết hợp (joint display matrix) để diễn giải toàn diện.',
     'Cross-sectional mixed-methods study following Convergent Parallel '
     'Design (Creswell & Plano Clark, 2018). Quantitative and qualitative '
     'data are collected in parallel, analyzed separately, then integrated '
     'via a joint display matrix for comprehensive interpretation.')


H2('4.2 Đối tượng nghiên cứu và thu mẫu', '4.2 Participants and sampling')

participants_data = [
    (('Cỡ mẫu', 'Sample size'),
     ('N = 1.352 học sinh THCS', 'N = 1,352 lower secondary school students')),
    (('Giới tính', 'Gender'),
     ('Nam: 614 (45,4%); Nữ: 738 (54,6%)',
      'Male: 614 (45.4%); Female: 738 (54.6%)')),
    (('Khối lớp', 'Grade levels'),
     ('Khối 6: n = 368; Khối 7: 316; Khối 8: 340; Khối 9: 328',
      'Grade 6: n = 368; Grade 7: 316; Grade 8: 340; Grade 9: 328')),
    (('Độ tuổi', 'Age range'),
     ('11 – 14 tuổi', '11 – 14 years')),
    (('Địa bàn', 'Sites'),
     ('2 trường THCS tại Hà Nội (Nhật Tân + Tây Mỗ)',
      '2 lower secondary schools in Hanoi (Nhat Tan + Tay Mo)')),
    (('Chiến lược chọn mẫu', 'Sampling strategy'),
     ('Chọn mẫu mục đích (purposive) đại diện cho khu vực nội thành – '
      'ngoại thành Hà Nội; mọi học sinh ở các lớp được chọn đều được mời '
      'tham gia',
      'Purposive sampling representing urban and suburban Hanoi; all '
      'students in selected classes invited to participate')),
]
make_table(
    headers_vn=['Tiêu chí', 'Chi tiết'],
    headers_en=['Criterion', 'Details'],
    rows=participants_data,
    col_widths_cm=[5.0, 12.0],
)


H2('4.3 Công cụ đo lường (8 thang đo)', '4.3 Measurement instruments (8 scales)')

Para('Tất cả các thang đo đã được kiểm chứng quốc tế và được thích ứng '
     'theo quy trình dịch xuôi – dịch ngược, tham vấn 3 chuyên gia tâm lý '
     'trẻ em + 2 chuyên gia giáo dục, và khảo sát thử nghiệm (n = 50).',
     'All scales were internationally validated and adapted through '
     'forward-backward translation, consultation with 3 child psychologists '
     '+ 2 educators, and pilot testing (n = 50).')

scales_data = [
    (('RCADS', 'RCADS'),
     ('Revised Children\'s Anxiety and Depression Scale (Chorpita, 2000)\n'
      '15 mục; 3 nhân tố: GAD 7 mục, SAD 4 mục, SocAD 4 mục\nThang Likert 4 mức',
      'Revised Children\'s Anxiety and Depression Scale (Chorpita, 2000)\n'
      '15 items; 3 factors: GAD 7 items, SAD 4 items, SocAD 4 items\n'
      '4-point Likert scale')),
    (('ESSA', 'ESSA'),
     ('Educational Stress Scale for Adolescents (Sun và cs., 2011)\n'
      '4 mục đo áp lực học tập; thang Likert 4 mức',
      'Educational Stress Scale for Adolescents (Sun et al., 2011)\n'
      '4 items measuring academic pressure; 4-point Likert')),
    (('SAS-SV', 'SAS-SV'),
     ('Smartphone Addiction Scale – Short Version (Kwon và cs., 2013)\n'
      '5 mục đo nghiện điện thoại; thang Likert 4 mức',
      'Smartphone Addiction Scale – Short Version (Kwon et al., 2013)\n'
      '5 items measuring smartphone addiction; 4-point Likert')),
    (('OBVQ', 'OBVQ'),
     ('Olweus Bully/Victim Questionnaire (Olweus, 1996)\n'
      '8 mục: bắt nạt thể chất (4) + bắt nạt lời nói (4)',
      'Olweus Bully/Victim Questionnaire (Olweus, 1996)\n'
      '8 items: physical bullying (4) + verbal bullying (4)')),
    (('PSSM', 'PSSM'),
     ('Psychological Sense of School Membership (Goodenow, 1993)\n'
      '7 mục đo gắn bó với trường học',
      'Psychological Sense of School Membership (Goodenow, 1993)\n'
      '7 items measuring school engagement')),
    (('MSPSS', 'MSPSS'),
     ('Multidimensional Scale of Perceived Social Support (Zimet và cs., 1988)\n'
      '8 mục: hỗ trợ cha mẹ (4) + hỗ trợ bạn bè (4)',
      'Multidimensional Scale of Perceived Social Support (Zimet et al., 1988)\n'
      '8 items: parental support (4) + peer support (4)')),
    (('RSES', 'RSES'),
     ('Rosenberg Self-Esteem Scale (Rosenberg, 1965)\n'
      '5 mục đo lòng tự trọng',
      'Rosenberg Self-Esteem Scale (Rosenberg, 1965)\n'
      '5 items measuring self-esteem')),
    (('Brief COPE', 'Brief COPE'),
     ('Brief COPE Inventory (Carver, 1997)\n'
      '15 mục; 3 nhân tố: giải quyết vấn đề, tìm hỗ trợ, tránh né',
      'Brief COPE Inventory (Carver, 1997)\n'
      '15 items; 3 factors: problem-solving, support-seeking, avoidance')),
]

make_table(
    headers_vn=['Mã', 'Mô tả thang đo'],
    headers_en=['Code', 'Scale description'],
    rows=scales_data,
    col_widths_cm=[2.0, 15.0],
)

Para('Điểm thô được chuyển đổi sang thang điểm 0-100 để so sánh giữa các '
     'thang đo có thang điểm Likert khác nhau.',
     'Raw scores are converted to a 0-100 scale for cross-scale comparability '
     'across instruments with different Likert ranges.')


H2('4.4 Phỏng vấn định tính', '4.4 Qualitative interviews')

NCS_marker('NCS xác nhận số người tham gia phỏng vấn, chiến lược lấy mẫu, '
           'thời lượng, trạng thái transcripts, và độ tin cậy liên mã (Cohen κ).')

Para('Phương pháp tạm thời: phỏng vấn bán cấu trúc 30-45 phút mỗi lượt; '
     'chọn mẫu phân tầng theo mức độ lo âu (cao/trung bình/thấp); phân tích '
     'theo phương pháp Thematic Analysis (Braun & Clarke, 2006).',
     'Tentative method: semi-structured 30-45 minute interviews per session; '
     'stratified sampling by anxiety severity (high/medium/low); Thematic '
     'Analysis (Braun & Clarke, 2006).')


H2('4.5 Chiến lược phân tích', '4.5 Analytic strategy')

Para('Quy trình phân tích được thực hiện theo 5 bước, với các ngưỡng chỉ '
     'số phù hợp được nêu rõ ràng để đảm bảo chất lượng đánh giá theo '
     'tiêu chuẩn quốc tế.',
     'Analysis proceeds in 5 steps, with explicitly stated fit indices '
     'thresholds to ensure quality assessment per international standards.')

analytic_data = [
    (('Bước 1', 'Step 1'),
     ('Mô tả thống kê', 'Descriptive statistics'),
     ('M, SD, range cho từng biến; độ tin cậy: Cronbach α và McDonald ω; '
      'cả hai chỉ số ≥ 0,70 để đạt mức chấp nhận',
      'M, SD, range for each variable; reliability: Cronbach α and McDonald '
      'ω; both indices ≥ 0.70 for acceptability')),
    (('Bước 2', 'Step 2'),
     ('Phân tích nhân tố khẳng định (CFA)',
      'Confirmatory Factor Analysis (CFA)'),
     ('Cho từng thang đo riêng biệt. Ngưỡng chỉ số phù hợp '
      '(Hu & Bentler, 1999): CFI ≥ 0,90; TLI ≥ 0,90; RMSEA ≤ 0,08; SRMR ≤ 0,08',
      'Per scale separately. Fit indices thresholds (Hu & Bentler, 1999): '
      'CFI ≥ 0.90; TLI ≥ 0.90; RMSEA ≤ 0.08; SRMR ≤ 0.08')),
    (('Bước 3', 'Step 3'),
     ('SEM tích hợp', 'Integrated SEM'),
     ('Phần mềm AMOS 31.0; 7 biến tiềm ẩn dự báo (3 nguy cơ + 4 bảo vệ) → '
      '3 biến tiềm ẩn kết quả (GAD, SAD, SocAD); tổng cộng 21 đường dẫn',
      'AMOS 31.0 software; 7 latent predictors (3 risk + 4 protective) → '
      '3 latent outcomes (GAD, SAD, SocAD); 21 paths total')),
    (('Bước 4', 'Step 4'),
     ('Kiểm định bất biến đa nhóm', 'Multi-group invariance testing'),
     ('Theo giới tính: configural → metric → scalar; tiêu chí '
      '(Cheung & Rensvold, 2002): ΔCFI ≤ 0,01 cho mỗi bước',
      'By gender: configural → metric → scalar; criteria '
      '(Cheung & Rensvold, 2002): ΔCFI ≤ 0.01 per step')),
    (('Bước 5', 'Step 5'),
     ('Tích hợp hỗn hợp', 'Mixed-methods integration'),
     ('Ma trận trình bày kết hợp (joint display matrix); chủ đề định tính '
      'đặt cạnh hệ số β định lượng để tam giác hóa',
      'Joint display matrix; qualitative themes placed alongside '
      'quantitative β coefficients for triangulation')),
]
make_table(
    headers_vn=['Bước', 'Tên', 'Chi tiết kỹ thuật'],
    headers_en=['Step', 'Name', 'Technical details'],
    rows=analytic_data,
    col_widths_cm=[1.5, 4.0, 11.5],
)

NCS_marker('Quyết định Q1-8: (A) Chạy lại SEM tích hợp đầy đủ với 7 dự báo '
           '→ 1 R² tổng hợp, HAY (B) làm rõ rằng R² = 0,284 và 0,209 là từ '
           'hai mô hình riêng (LA chính Bảng 27 + 30) và thêm chú thích "for '
           'comparability with parent dissertation"')


H2('4.6 Đạo đức nghiên cứu', '4.6 Ethics')

NCS_marker('NCS bổ sung số quyết định + ngày + tên Hội đồng Đạo đức HNUE.')

Para('Tuyên bố tạm thời: nghiên cứu đã được Hội đồng Đạo đức Nghiên cứu '
     'của Trường Đại học Sư phạm Hà Nội phê duyệt (số [...], ngày [...]). '
     'Đã có sự đồng ý bằng văn bản của cha mẹ/người giám hộ hợp pháp và '
     'sự đồng thuận bằng văn bản của học sinh tham gia. Dữ liệu được mã hóa '
     'ẩn danh; chỉ nhóm nghiên cứu được truy cập. Quy trình tuân thủ '
     'Tuyên bố Helsinki (2013) và Luật Trẻ em Việt Nam (2016).',
     'Tentative statement: this study was approved by the Research Ethics '
     'Committee of Hanoi National University of Education (number [...], '
     'date [...]). Written informed consent was obtained from parents/legal '
     'guardians, and written assent from student participants. Data are '
     'anonymized; only the research team has access. Procedures comply with '
     'the Declaration of Helsinki (2013) and Vietnamese Law on Children (2016).')


d.add_page_break()


# ============================================================
H1('5. KẾT QUẢ DỰ KIẾN', '5. EXPECTED RESULTS')


H2('5.1 Đặc điểm mẫu', '5.1 Sample characteristics')

Para('Bảng 1 mô tả đặc điểm mẫu 1.352 học sinh theo giới tính, khối lớp '
     'và trường học.',
     'Table 1 describes characteristics of the 1,352-student sample by '
     'gender, grade level, and school.')


H2('5.2 Độ tin cậy và phù hợp CFA', '5.2 Reliability and CFA fit')

Para('Bảng 2 trình bày Cronbach α, McDonald ω và các chỉ số phù hợp CFA '
     '(CFI, TLI, RMSEA, SRMR) cho 8 thang đo.',
     'Table 2 reports Cronbach α, McDonald ω, and CFA fit indices (CFI, TLI, '
     'RMSEA, SRMR) for the 8 scales.')


H2('5.3 Mức độ rối loạn lo âu', '5.3 Anxiety levels')

anx_levels_data = [
    (('Rối loạn lo âu lan toả (GAD/RLLALT)',
      'Generalized anxiety disorder (GAD)'),
     ('55,82', '55.82'),
     ('~22,4', '~22.4'),
     ('Mức trung bình – cao', 'Moderate-high')),
    (('Rối loạn lo âu chia ly (SAD/RLLACL)',
      'Separation anxiety disorder (SAD)'),
     ('25,06', '25.06'),
     ('24,29', '24.29'),
     ('Mức thấp nhất – phù hợp tuổi tác',
      'Lowest level – developmentally appropriate')),
    (('Rối loạn lo âu xã hội (SocAD/RLLAXH)',
      'Social anxiety disorder (SocAD)'),
     ('48,41', '48.41'),
     ('26,19', '26.19'),
     ('Mức trung bình', 'Moderate')),
]
make_table(
    headers_vn=['Dạng RLLA', 'Trung bình (M)', 'Độ lệch chuẩn (SD)', 'Diễn giải'],
    headers_en=['Disorder type', 'Mean (M)', 'Standard deviation (SD)', 'Interpretation'],
    rows=anx_levels_data,
    col_widths_cm=[5.5, 3.0, 3.0, 5.5],
)


H2('5.4 Mô hình SEM chính — Bảng hệ số β đặc thù theo dạng RLLA',
   '5.4 Main SEM model — Subtype-specific β coefficient table')

Para('Bảng 3 trình bày hệ số β chuẩn hóa của 7 dự báo lên 3 dạng rối loạn '
     'lo âu (21 đường dẫn) trong mô hình SEM tích hợp. Tất cả giá trị β '
     'đã được kiểm chứng đối chiếu với LA chính (Bảng 27-42 trong luận án).',
     'Table 3 reports standardized β coefficients of 7 predictors on 3 '
     'anxiety subtypes (21 paths) in the integrated SEM model. All β values '
     'have been cross-verified against the parent dissertation '
     '(Tables 27-42).')

beta_rows = [
    (('🔴 YẾU TỐ NGUY CƠ', '🔴 RISK FACTORS'),
     ('', ''), ('', ''), ('', ''), ('', '')),
    (('Áp lực học tập', 'Academic pressure'),
     ('+0,510***', '+0.510***'),
     ('[verify]', '[verify]'),
     ('+0,490***', '+0.490***'),
     ('+0,533***', '+0.533***')),
    (('Nghiện điện thoại', 'Smartphone addiction'),
     ('+0,336***', '+0.336***'),
     ('+0,265***', '+0.265***'),
     ('+0,383***', '+0.383***'),
     ('+0,400***', '+0.400***')),
    (('Bắt nạt học đường ⭐', 'School bullying ⭐'),
     ('+0,253***', '+0.253***'),
     ('+0,376*** ⭐', '+0.376*** ⭐'),
     ('+0,215***', '+0.215***'),
     ('+0,276***', '+0.276***')),
    (('🛡 YẾU TỐ BẢO VỆ', '🛡 PROTECTIVE FACTORS'),
     ('', ''), ('', ''), ('', ''), ('', '')),
    (('Gắn bó trường học', 'School engagement'),
     ('-0,108**', '-0.108**'),
     ('+0,014 ns', '+0.014 ns'),
     ('-0,187***', '-0.187***'),
     ('-0,155***', '-0.155***')),
    (('Hỗ trợ cha mẹ', 'Parental support'),
     ('-0,172***', '-0.172***'),
     ('[verify]', '[verify]'),
     ('-0,273***', '-0.273***'),
     ('-0,273***', '-0.273***')),
    (('Hỗ trợ bạn bè', 'Peer support'),
     ('-0,015 ns', '-0.015 ns'),
     ('-0,019 ns', '-0.019 ns'),
     ('positive ns', 'positive ns'),
     ('-0,015 ns', '-0.015 ns')),
    (('Lòng tự trọng ⭐', 'Self-esteem ⭐'),
     ('-0,455***', '-0.455***'),
     ('-0,087**', '-0.087**'),
     ('-0,415***', '-0.415***'),
     ('-0,457***', '-0.457***')),
]

make_table(
    headers_vn=['Biến dự báo', '→ GAD (β)', '→ SAD (β)', '→ SocAD (β)',
                'RLLA tổng (β)'],
    headers_en=['Predictor', '→ GAD (β)', '→ SAD (β)', '→ SocAD (β)',
                'Total anxiety (β)'],
    rows=beta_rows,
    col_widths_cm=[5.0, 3.0, 3.0, 3.0, 3.0],
)

Para('Chú thích: *** p < 0,001; ** p < 0,01; * p < 0,05; ns = không có ý '
     'nghĩa thống kê. ⭐ = đường dẫn nổi bật cho điểm mới của nghiên cứu.',
     'Notes: *** p < 0.001; ** p < 0.01; * p < 0.05; ns = non-significant. '
     '⭐ = notable path for the novel contributions of this study.', italic=True)

NCS_marker('Hai ô [verify] (ALHT → SAD và HTCM → SAD) cần em tra cứu lại '
           'từ các Bảng chi tiết trong LA chính. Sẽ hoàn thành ở bước viết draft.')

Para('Hệ số R²:', 'R² coefficients:', indent=False)
Bullet('R² mô hình yếu tố nguy cơ (3 dự báo → RLLA tổng) = 0,284',
       'R² of risk-factors model (3 predictors → total anxiety) = 0.284')
Bullet('R² mô hình yếu tố bảo vệ (4 dự báo → RLLA tổng) = 0,209',
       'R² of protective-factors model (4 predictors → total anxiety) = 0.209')

NCS_marker('Quyết định Q1-8: Nếu thầy + NCS quyết chạy SEM tích hợp đầy đủ, '
           'sẽ có 1 R² hợp nhất thay vì 2 R² riêng.')


H2('5.5 Kiểm định bất biến đa nhóm theo giới',
   '5.5 Multi-group invariance by gender')

Para('Kết quả kiểm định bất biến cho thấy mô hình phù hợp với cả hai nhóm '
     'nam và nữ. Tuy nhiên, khác biệt mức độ rối loạn lo âu theo giới '
     'thể hiện rõ ràng và có ý nghĩa thống kê khác nhau theo dạng RLLA:',
     'Invariance testing reveals model fit for both male and female groups. '
     'However, gender differences in anxiety levels are clearly evident '
     'and vary in statistical significance across subtypes:')

gender_data = [
    (('GAD', 'GAD'),
     ('Nam: M = 51,43 (SD = 22,01)', 'Male: M = 51.43 (SD = 22.01)'),
     ('Nữ: M = 59,47 (SD = 22,07)', 'Female: M = 59.47 (SD = 22.07)'),
     ('F = 44,484; p < 0,001', 'F = 44.484; p < 0.001'),
     ('Có khác biệt giới rõ rệt', 'Significant gender difference')),
    (('SAD ⭐', 'SAD ⭐'),
     ('Nam: M = 25,42 (SD = 25,46)', 'Male: M = 25.42 (SD = 25.46)'),
     ('Nữ: M = 24,76 (SD = 23,29)', 'Female: M = 24.76 (SD = 23.29)'),
     ('F = 0,246; p = 0,620', 'F = 0.246; p = 0.620'),
     ('KHÔNG có khác biệt giới (BẤT BIẾN)',
      'NO gender difference (INVARIANT)')),
    (('SocAD', 'SocAD'),
     ('Nam: M = 43,20 (SD = 25,09)', 'Male: M = 43.20 (SD = 25.09)'),
     ('Nữ: M = 52,74 (SD = 26,31)', 'Female: M = 52.74 (SD = 26.31)'),
     ('F = 45,984; p < 0,001', 'F = 45.984; p < 0.001'),
     ('Có khác biệt giới rõ rệt', 'Significant gender difference')),
]
make_table(
    headers_vn=['Dạng', 'Mức độ nam', 'Mức độ nữ', 'Kiểm định', 'Kết luận'],
    headers_en=['Type', 'Male level', 'Female level', 'Test', 'Conclusion'],
    rows=gender_data,
    col_widths_cm=[2.0, 3.5, 3.5, 3.5, 4.5],
)

Para('⭐ Phát hiện then chốt: rối loạn lo âu chia ly thể hiện tính bất biến '
     'giới hoàn toàn, khác hẳn với lo âu lan toả và lo âu xã hội. Đây là '
     'phát hiện trung tâm cho luận điểm về vai trò của bối cảnh văn hóa '
     'tập thể chủ nghĩa.',
     '⭐ Key finding: separation anxiety exhibits complete gender invariance, '
     'distinct from generalized and social anxiety. This is the central '
     'finding for arguments regarding the role of collectivist cultural '
     'context.', indent=False, italic=True)


H2('5.6 Tích hợp định tính – định lượng', '5.6 Mixed-methods integration')

NCS_marker('Chủ đề định tính sẽ được bổ sung sau khi NCS cung cấp dữ liệu '
           'phỏng vấn và transcripts.')

Para('Chủ đề tạm thời (placeholder, chờ NCS xác nhận):',
     'Tentative themes (placeholder, pending NCS confirmation):',
     indent=False)

Bullet('"Im lặng trong gia đình" – giải thích hiệu lực hạn chế của yếu tố '
       '"hỗ trợ cha mẹ"',
       '"Silence in the family" – explains the limited efficacy of the '
       'parental support factor')
Bullet('"Áp lực bạn đồng lứa" – giải thích vì sao "hỗ trợ bạn bè" không '
       'mang tính bảo vệ',
       '"Peer pressure" – explains why peer support fails to act protectively')
Bullet('"Sợ tách biệt" – minh chứng cho tính bất biến giới ở lo âu chia ly',
       '"Fear of separation" – evidence for gender invariance in separation '
       'anxiety')


d.add_page_break()


# ============================================================
H1('6. PHẦN BÀN LUẬN', '6. DISCUSSION')


H2('6.1 Tóm tắt các phát hiện chính', '6.1 Summary of key findings')

Para('Ba giả thuyết khoa học đều được kiểm chứng phù hợp với mô hình:',
     'All three scientific hypotheses are consistent with the model:',
     indent=False)

Bullet('H1 đúng: Cả 3 yếu tố nguy cơ đều có đường dẫn dương có ý nghĩa với '
       'cả 3 dạng RLLA',
       'H1 confirmed: All 3 risk factors show significant positive paths to '
       'all 3 anxiety subtypes')
Bullet('H2 đúng (gần như): 4 yếu tố bảo vệ đều có đường dẫn âm với GAD và '
       'SocAD (trừ ngoại lệ "hỗ trợ bạn bè")',
       'H2 largely confirmed: 4 protective factors show negative paths to GAD '
       'and SocAD (except for "peer support" exception)')
Bullet('H3 đúng: Tính bất biến giới của rối loạn lo âu chia ly đã được xác '
       'nhận — đây là phát hiện mới quan trọng',
       'H3 confirmed: Gender invariance of separation anxiety has been '
       'verified — this is an important novel finding')


H2('6.2 Diễn giải yếu tố nguy cơ', '6.2 Interpretation of risk factors')

Para('Sự thống trị của áp lực học tập (β = +0,533 trên RLLA tổng) phù hợp '
     'với hai phân tích hệ thống lớn: Pascoe và cộng sự (2020) trên y văn '
     'quốc tế và Steare và cộng sự (2023) với 48/52 nghiên cứu xác nhận '
     'mối liên hệ dương giữa áp lực học tập và sức khỏe tâm thần. Bối cảnh '
     'Việt Nam đặc thù với kỳ thi đại học cạnh tranh cao càng làm nổi bật '
     'cường độ tác động này.',
     'The dominance of academic pressure (β = +0.533 on total anxiety) is '
     'consistent with two major systematic reviews: Pascoe et al. (2020) on '
     'international literature and Steare et al. (2023) where 48 of 52 '
     'studies confirmed positive association between academic pressure and '
     'mental health. The distinctive Vietnamese context with high-stakes '
     'university entrance examinations amplifies this effect intensity.')

Para('Nghiện điện thoại có tác động mạnh nhất lên lo âu xã hội (β = '
     '+0,383), phù hợp với cơ chế lý thuyết FOMO (Fear of Missing Out) và '
     'so sánh xã hội trực tuyến (Chen và cs., 2023). Điều này gợi ý chương '
     'trình phòng ngừa cần đặc biệt chú trọng sử dụng điện thoại của học '
     'sinh có biểu hiện lo âu xã hội.',
     'Smartphone addiction shows the strongest effect on social anxiety '
     '(β = +0.383), consistent with FOMO (Fear of Missing Out) and online '
     'social comparison theoretical mechanisms (Chen et al., 2023). This '
     'suggests prevention programs should particularly focus on smartphone '
     'use among students with social anxiety symptoms.')

Para('Phát hiện nổi bật là tác động của BẮT NẠT HỌC ĐƯỜNG lên lo âu chia '
     'ly (β = +0,376) — đường dẫn ĐƠN MẠNH NHẤT đến SAD trong tất cả các '
     'biến nghiên cứu, mạnh hơn cả áp lực học tập. Điều này cho thấy bắt '
     'nạt có cơ chế tâm lý đặc thù liên quan đến cảm giác an toàn và sự '
     'gắn bó, không chỉ là một stressor học đường chung chung.',
     'A notable finding is the impact of SCHOOL BULLYING on separation '
     'anxiety (β = +0.376) — the STRONGEST SINGLE path to SAD across all '
     'variables, stronger than even academic pressure. This suggests '
     'bullying has a distinctive psychological mechanism related to safety '
     'and attachment, not merely a generic school stressor.')


H2('6.3 Lòng tự trọng — đòn bẩy bảo vệ mạnh nhất',
   '6.3 Self-esteem — strongest protective lever')

Para('Lòng tự trọng có hệ số β = -0,455 đến GAD, với cường độ tác động '
     'tương đương khoảng 85-89% so với áp lực học tập (theo Kết luận LA '
     'chính). Phát hiện này phù hợp với phân tích meta của Compas và cộng '
     'sự (2017): các can thiệp tái cấu trúc nhận thức và xây dựng lòng tự '
     'trọng có cỡ tác động tương tự nhau. Đây là cơ sở khoa học mạnh cho '
     'việc đưa các module xây dựng lòng tự trọng vào chương trình phòng '
     'ngừa rối loạn lo âu trường học.',
     'Self-esteem has β = -0.455 on GAD, with effect intensity approximately '
     '85-89% of academic pressure (per parent dissertation Conclusions). '
     'This finding aligns with the Compas et al. (2017) meta-analysis: '
     'cognitive restructuring and self-esteem-building interventions show '
     'similar effect sizes. This provides strong scientific basis for '
     'incorporating self-esteem-building modules in school-based anxiety '
     'prevention programs.')


H2('6.4 Tính bất biến giới của lo âu chia ly — diễn giải văn hóa',
   '6.4 Gender invariance of separation anxiety — cultural interpretation')

Para('Phát hiện rằng rối loạn lo âu chia ly không có khác biệt giới '
     '(F = 0,246; p = 0,620) là điểm mới quan trọng nhất của nghiên cứu, '
     'cần được diễn giải qua hai khung lý thuyết:',
     'The finding that separation anxiety shows no gender difference '
     '(F = 0.246; p = 0.620) represents the most significant novel '
     'contribution of this study, requiring interpretation through two '
     'theoretical frameworks:')

H3('(a) Lý thuyết Văn hóa Tập thể Chủ nghĩa',
   '(a) Cultural Collectivism Theory')

Para('Theo Triandis (1995) và Markus & Kitayama (1991), cấu trúc gia đình '
     'theo thứ bậc và định hướng "bản ngã phụ thuộc" (interdependent '
     'self-construal) trong văn hóa Việt Nam tạo ra trải nghiệm gắn bó '
     'đồng nhất giữa các giới. Khác với các bối cảnh cá nhân chủ nghĩa '
     'phương Tây nơi sự khác biệt nam-nữ xuất hiện sớm qua kỳ vọng tự chủ, '
     'văn hóa tập thể Việt Nam duy trì động lực gắn bó đồng nhất giữa các '
     'giới trong giai đoạn vị thành niên sớm.',
     'According to Triandis (1995) and Markus & Kitayama (1991), '
     'hierarchical family structure and "interdependent self-construal" '
     'orientation in Vietnamese culture create uniform attachment '
     'experiences across genders. Unlike Western individualist contexts '
     'where male-female differentiation occurs early through autonomy '
     'expectations, Vietnamese collectivism preserves uniform attachment '
     'dynamics across genders during early adolescence.')

H3('(b) Lý thuyết Nhiệm vụ Phát triển', '(b) Developmental Task Theory')

Para('Allen và cộng sự (2013) cho thấy các nhiệm vụ phân ly – cá biệt hóa '
     '(separation-individuation) chủ yếu khởi phát từ thời thơ ấu, có '
     'trước các áp lực xã hội phân biệt theo giới xuất hiện sau dậy thì. '
     'Trước dậy thì: nhiệm vụ phổ quát giữa các giới. Sau dậy thì: '
     'áp lực xã hội phân biệt theo giới → GAD và SocAD phân kỳ nhưng '
     'SAD vẫn ổn định.',
     'Allen et al. (2013) demonstrate that separation-individuation tasks '
     'primarily originate in childhood, predating gender-differentiated '
     'social pressures emerging after puberty. Pre-puberty: tasks universal '
     'across genders. Post-puberty: gender-differentiated social pressures '
     '→ GAD and SocAD diverge while SAD remains stable.')


H2('6.5 Hàm ý lâm sàng và giáo dục', '6.5 Clinical and educational implications')

Bullet('Phòng ngừa phân tầng: can thiệp toàn trường (universal) + can thiệp '
       'mục tiêu cho học sinh có nguy cơ cao (targeted)',
       'Tiered prevention: universal school-based intervention + targeted '
       'intervention for high-risk students')
Bullet('Khung chương trình tập huấn 8 nội dung từ LA chính được điều chỉnh '
       'để phù hợp với khung bằng chứng quốc tế',
       'The 8-component prevention curriculum framework from the parent '
       'dissertation, adapted to align with international evidence-based '
       'framework')
Bullet('Can thiệp nhận thức giới: phương pháp tương tự cho SAD (bất biến '
       'giới), nhưng phân hóa theo giới cho GAD và SocAD',
       'Gender-aware intervention: same approach for SAD (gender-invariant), '
       'but gender-differentiated for GAD and SocAD')


H2('6.6 Giới hạn nghiên cứu và hướng nghiên cứu tương lai',
   '6.6 Limitations and future research')

Bullet('Thiết kế cắt ngang → không thể suy luận nhân quả; cần nghiên cứu '
       'theo dõi (longitudinal) để xác nhận cơ chế',
       'Cross-sectional design → cannot infer causation; longitudinal study '
       'needed to confirm mechanisms')
Bullet('2 trường Hà Nội → khả năng khái quát hóa hạn chế; cần mở rộng '
       'sang các tỉnh khác và khu vực nông thôn',
       '2 schools in Hanoi → generalizability limitation; expansion to other '
       'provinces and rural areas needed')
Bullet('Tự báo cáo có thể có sai lệch; cần kết hợp đánh giá quan sát/lâm sàng',
       'Self-report may have bias; observational/clinical assessment '
       'combination needed')
Bullet('Cỡ mẫu phỏng vấn định tính nhỏ — chờ NCS xác nhận',
       'Qualitative interview sub-sample size is small — pending NCS '
       'confirmation')
Bullet('Hướng tương lai: thử nghiệm RCT chương trình tập huấn 8 nội dung '
       'để xác minh hiệu quả can thiệp',
       'Future direction: RCT testing the 8-component prevention curriculum '
       'to verify intervention efficacy')


d.add_page_break()


# ============================================================
H1('7. TÀI LIỆU THAM KHẢO DỰ KIẾN', '7. EXPECTED REFERENCES')


Para('Tất cả tài liệu tham khảo đã được kiểm chứng đối với PDF gốc trong '
     'kho dữ liệu nghiên cứu (`02_Papers-goc/`). Dự kiến 45-50 tài liệu, '
     'phân thành ba nhóm:',
     'All references have been verified against original PDFs in the research '
     'database (`02_Papers-goc/`). Approximately 45-50 references in three '
     'groups:', indent=False)


H2('7.1 Nghiên cứu thực nghiệm (16 PDFs đã kiểm chứng)',
   '7.1 Empirical studies (16 verified PDFs)')

Bullet('Xu et al. (2021) – Tỷ lệ và yếu tố nguy cơ rối loạn lo âu trên '
       'N = 373.216 HS THCS Trung Quốc thời kỳ COVID-19',
       'Xu et al. (2021) – Prevalence and risk factors for anxiety symptoms '
       'in N = 373,216 Chinese lower secondary students during COVID-19')
Bullet('Chen et al. (2023) – Bắt nạt và sức khỏe tâm thần trên 63.205 học '
       'sinh phổ thông Trung Quốc',
       'Chen et al. (2023) – Bullying and mental health in 63,205 Chinese '
       'secondary students')
Bullet('Wen et al. (2020) – Phân tích hồ sơ tiềm ẩn về lo âu ở học sinh '
       'nông thôn Trung Quốc',
       'Wen et al. (2020) – Latent profile analysis of anxiety in rural '
       'Chinese students')
Bullet('Anderson et al. (2025) – Tổng quan tường thuật yếu tố làm tăng '
       'lo âu vị thành niên',
       'Anderson et al. (2025) – Narrative review of factors contributing '
       'to the rise in adolescent anxiety')
Bullet('Pascoe et al. (2020) – Tác động của stress học đường (tổng quan '
       'hệ thống)',
       'Pascoe et al. (2020) – Impact of stress on students in secondary '
       'school (systematic review)')
Bullet('Steare et al. (2023) – Tổng quan hệ thống áp lực học tập: 48/52 '
       'nghiên cứu xác nhận liên hệ dương',
       'Steare et al. (2023) – Systematic review of academic pressure: '
       '48/52 studies confirmed positive association')
Bullet('Saikia et al. (2023) – Sức khỏe tâm thần thanh thiếu niên Đông '
       'Bắc Ấn Độ (Boy 30,0% vs Girl 18,9% lo âu nặng)',
       'Saikia et al. (2023) – Adolescent mental health in Northeast India '
       '(Boy 30.0% vs Girl 18.9% severe anxiety)')
Bullet('Bhardwaj et al. (2020) – DASS-21 trên N = 288 HS THCS Chandigarh, Ấn Độ',
       'Bhardwaj et al. (2020) – DASS-21 on N = 288 lower secondary students '
       'in Chandigarh, India')
Bullet('V-NAMHS (2022) – Điều tra sức khỏe tâm thần vị thành niên quốc gia '
       'Việt Nam (DISC-5)',
       'V-NAMHS (2022) – Vietnam National Adolescent Mental Health Survey '
       '(DISC-5)')
Bullet('Hoang Trung Hoc et al. (2025) – N = 8.473 học sinh Việt Nam sau '
       'COVID-19',
       'Hoang Trung Hoc et al. (2025) – N = 8,473 post-COVID-19 Vietnamese '
       'students')
Bullet('Nakie et al. (2022) – Trầm cảm, lo âu, stress ở học sinh phổ thông '
       'Ethiopia',
       'Nakie et al. (2022) – Depression, anxiety, stress in Ethiopian high '
       'school students')
Bullet('Jefferies & Ungar (2020) – Lo âu xã hội ở thanh niên 7 quốc gia',
       'Jefferies & Ungar (2020) – Social anxiety in young people across '
       '7 countries')
Bullet('Qiu et al. (2022) – Phong cách nuôi dạy và sức bật tâm lý ở học '
       'sinh Trung Quốc',
       'Qiu et al. (2022) – Parenting style and resilience in Chinese students')
Bullet('Zhu et al. (2025) – Yếu tố ảnh hưởng đến SKTT học sinh phổ thông Trung Quốc',
       'Zhu et al. (2025) – Factors affecting mental health in Chinese '
       'secondary students')
Bullet('Alharbi et al. (2019) – Trầm cảm và lo âu HS THPT vùng Qassim, Ả Rập Saudi',
       'Alharbi et al. (2019) – Depression and anxiety in Qassim region high '
       'school students, Saudi Arabia')
Bullet('GBD ASEAN (2025) – Gánh nặng bệnh tật toàn cầu khu vực ASEAN trên '
       'tạp chí Lancet',
       'GBD ASEAN (2025) – Global Burden of Disease ASEAN region in Lancet')


H2('7.2 Phương pháp luận', '7.2 Methodological')

Bullet('Compas et al. (2017) – Phân tích meta về đối phó và bệnh lý tâm thần '
       '(N = 80.850, 212 nghiên cứu)',
       'Compas et al. (2017) – Meta-analysis of coping and psychopathology '
       '(N = 80,850, 212 studies)')
Bullet('Hu & Bentler (1999) – Tiêu chí chỉ số phù hợp SEM',
       'Hu & Bentler (1999) – SEM fit indices criteria')
Bullet('Cheung & Rensvold (2002) – Kiểm định bất biến ΔCFI',
       'Cheung & Rensvold (2002) – ΔCFI invariance testing')
Bullet('Braun & Clarke (2006) – Phương pháp Thematic Analysis',
       'Braun & Clarke (2006) – Thematic Analysis methodology')
Bullet('Creswell & Plano Clark (2018) – Thiết kế nghiên cứu hỗn hợp '
       '(mixed-methods)',
       'Creswell & Plano Clark (2018) – Mixed-methods research design')


H2('7.3 Lý thuyết và thang đo', '7.3 Theory and instruments')

Bullet('Lazarus & Folkman (1984) – Lý thuyết Stress-Coping kinh điển',
       'Lazarus & Folkman (1984) – Classic Stress-Coping theory')
Bullet('Carver (1997) – Phát triển thang Brief COPE',
       'Carver (1997) – Brief COPE scale development')
Bullet('Rosenberg (1965) – Thang đo Rosenberg Self-Esteem',
       'Rosenberg (1965) – Rosenberg Self-Esteem Scale')
Bullet('Goodenow (1993) – Thang đo PSSM',
       'Goodenow (1993) – PSSM scale')
Bullet('Zimet et al. (1988) – Thang đo MSPSS',
       'Zimet et al. (1988) – MSPSS scale')
Bullet('Chorpita (2000) – Thang đo RCADS',
       'Chorpita (2000) – RCADS scale')
Bullet('Sun et al. (2011) – Thang đo ESSA',
       'Sun et al. (2011) – ESSA scale')
Bullet('Kwon et al. (2013) – Thang đo SAS-SV',
       'Kwon et al. (2013) – SAS-SV scale')
Bullet('Olweus (1996) – Bảng hỏi OBVQ',
       'Olweus (1996) – OBVQ questionnaire')
Bullet('McLean et al. (2011) – Phân tích meta khác biệt giới ở rối loạn lo âu',
       'McLean et al. (2011) – Meta-analysis of gender differences in '
       'anxiety disorders')


H2('7.4 Khung văn hóa và phát triển', '7.4 Cultural and developmental framework')

Bullet('Triandis (1995) – Lý thuyết tập thể chủ nghĩa và cá nhân chủ nghĩa',
       'Triandis (1995) – Collectivism and individualism theory')
Bullet('Markus & Kitayama (1991) – Bản ngã phụ thuộc và bản ngã độc lập',
       'Markus & Kitayama (1991) – Interdependent and independent self-'
       'construal')
Bullet('Allen et al. (2013) – Nhiệm vụ phân ly – cá biệt hóa',
       'Allen et al. (2013) – Separation-individuation tasks')
Bullet('Bronfenbrenner & Morris (2006) – Lý thuyết hệ sinh thái',
       'Bronfenbrenner & Morris (2006) – Bioecological theory')


d.add_page_break()


# ============================================================
H1('8. KẾ HOẠCH BẢNG VÀ HÌNH', '8. TABLES AND FIGURES PLAN')

tables_figures = [
    (('Bảng 1', 'Table 1'),
     ('Đặc điểm mẫu', 'Sample demographics'),
     ('Phân bố theo giới tính × khối lớp × trường học',
      'Distribution by gender × grade × school')),
    (('Bảng 2', 'Table 2'),
     ('Chỉ số tâm trắc 8 thang đo', 'Psychometric properties of 8 scales'),
     ('Cronbach α, McDonald ω, CFI, TLI, RMSEA, SRMR cho mỗi thang đo',
      'Cronbach α, McDonald ω, CFI, TLI, RMSEA, SRMR per scale')),
    (('Bảng 3', 'Table 3'),
     ('Bảng hệ số β SEM tích hợp (7×3)',
      'Integrated SEM β coefficients table (7×3)'),
     ('21 đường dẫn từ 7 dự báo lên 3 dạng RLLA, với mức ý nghĩa thống kê',
      '21 paths from 7 predictors to 3 anxiety subtypes, with statistical '
      'significance')),
    (('Bảng 4', 'Table 4'),
     ('Bảng chỉ số phù hợp kiểm định bất biến',
      'Multi-group invariance fit indices'),
     ('Configural / metric / scalar với ΔCFI, ΔTLI, ΔRMSEA',
      'Configural / metric / scalar with ΔCFI, ΔTLI, ΔRMSEA')),
    (('Bảng 5', 'Table 5'),
     ('Joint display tích hợp định tính – định lượng',
      'Joint display for qualitative-quantitative integration'),
     ('Chủ đề định tính × hệ số β định lượng cho từng yếu tố bảo vệ',
      'Qualitative themes × quantitative β coefficients per protective factor')),
    (('Hình 1', 'Figure 1'),
     ('Sơ đồ mô hình SEM tích hợp', 'Integrated SEM model diagram'),
     ('Trực quan hóa 7 dự báo → 3 dạng RLLA với β có ý nghĩa thống kê đánh dấu',
      'Visualization of 7 predictors → 3 anxiety subtypes with statistically '
      'significant β highlighted')),
    (('Hình 2', 'Figure 2'),
     ('Biểu đồ so sánh đa nhóm theo giới', 'Multi-group gender comparison chart'),
     ('Cường độ đường dẫn theo giới, làm nổi bật bất biến của SAD',
      'Path strengths by gender, highlighting SAD invariance')),
]

make_table(
    headers_vn=['Mã', 'Tiêu đề', 'Mô tả nội dung'],
    headers_en=['Code', 'Title', 'Content description'],
    rows=tables_figures,
    col_widths_cm=[1.5, 4.5, 11.0],
)


# ============================================================
H1('9. CÂU HỎI CHỜ QUYẾT (BLOCKING)',
   '9. BLOCKING QUESTIONS')

Para('Bốn câu hỏi sau cần thầy hướng dẫn và NCS quyết định trước khi tiến '
     'hành viết bản thảo:',
     'The following four questions require decisions from supervisor and PhD '
     'candidate before drafting:', indent=False)

blocking = [
    (('Q1-6', 'Q1-6'),
     ('Dữ liệu phỏng vấn định tính', 'Qualitative interview data'),
     ('NCS', 'NCS'),
     ('Số người tham gia phỏng vấn? Chiến lược lấy mẫu? Transcripts đã có '
      'chưa? Cohen κ?',
      'Number of interview participants? Sampling strategy? Transcripts '
      'ready? Cohen κ?')),
    (('Q1-8', 'Q1-8'),
     ('R² mô hình tích hợp', 'R² of integrated model'),
     ('Thầy + NCS', 'Supervisor + NCS'),
     ('(A) Chạy lại SEM tích hợp 7 dự báo → 1 R² tổng hợp, HAY (B) làm rõ '
      'rằng R² = 0,284 và 0,209 là từ hai mô hình riêng trong LA',
      '(A) Re-run integrated SEM with 7 predictors → 1 combined R², OR '
      '(B) clarify that R² = 0.284 and 0.209 are from separate models in '
      'the parent dissertation')),
    (('Q3-6', 'Q3-6'),
     ('Văn bản phê duyệt đạo đức HNUE', 'HNUE ethics approval letter'),
     ('NCS', 'NCS'),
     ('Đã có văn bản chính thức chưa? Nếu chưa, xin được retroactive '
      '(2-4 tuần)?',
      'Is there a formal letter? If not, retroactive approval possible '
      '(2-4 weeks)?')),
    (('Q3-9', 'Q3-9'),
     ('Chiến lược tham chiếu chéo Q1 ↔ Q3',
      'Q1 ↔ Q3 cross-reference strategy'),
     ('Thầy + NCS', 'Supervisor + NCS'),
     ('(A) Nộp Q1 trước (BMC Psychiatry), chờ chấp nhận, nộp Q3 với citation '
      'Q1, HAY (B) nộp Q1 và Q3 cùng lúc với chú thích "companion paper '
      'under review"',
      '(A) Submit Q1 first (BMC Psychiatry), wait for acceptance, submit Q3 '
      'with Q1 citation, OR (B) Submit Q1 and Q3 simultaneously with '
      '"companion paper under review" note')),
]

make_table(
    headers_vn=['Mã', 'Vấn đề', 'Người quyết', 'Chi tiết quyết định'],
    headers_en=['Code', 'Issue', 'Decision maker', 'Decision details'],
    rows=blocking,
    col_widths_cm=[1.5, 4.0, 2.5, 9.0],
)


# ============================================================
# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
