# -*- coding: utf-8 -*-
"""SUA PARA 223 LA CHINH - cac loi VN papers chua duoc sua truoc do.
29/05/2026.

LOI:
- Bao Quyen: 1.252 -> 501; 38,5% -> 86,2%
- Lam 2022: 1.024 -> 482; ban do thi -> ban nong thon; 25,8% -> 49,0%
- An Giang: Le Minh T. -> Nguyen Dang Khoa va cs.; 600 -> 366; 33,2% -> 61,2%
- Yen Dinh dia ly: ban do thi -> huyen ban nong thon
"""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

BACKUP = FILE.replace('.docx', '_BEFORE_FIX_PARA223.docx')
shutil.copy(FILE, BACKUP)
print(f"Backup: {BACKUP}")

RED = RGBColor(192, 0, 0)

REPLACEMENTS = [
    # Bao Quyen
    ('Nguyễn Ngọc Bảo Quyên và cộng sự (2025) khảo sát trên 1.252 học sinh trung học phổ thông Hà Nội bằng thang DASS-21 phiên bản tiếng Việt báo cáo tỷ lệ lo âu chung 38,5%, với nữ cao hơn nam có ý nghĩa thống kê.',
     'Nguyễn Ngọc Bảo Quyên và cộng sự (2025) khảo sát trên 501 học sinh trung học phổ thông Hà Nội bằng thang DASS-21 phiên bản tiếng Việt báo cáo tỷ lệ lo âu chung 86,2% (phân tầng: nhẹ 6,6%; vừa 25,1%; nặng 17,4%; rất nặng 37,1%), với nữ cao hơn nam có ý nghĩa thống kê (p < 0,001).'),
    # Lam 2022
    ('Nguyễn Danh Lâm và cộng sự (2022) sử dụng thang đo DASS-21 khảo sát trên 1.024 học sinh trung học phổ thông tại Yên Định, Thanh Hóa - khu vực bán đô thị - báo cáo tỷ lệ thấp hơn 25,8%.',
     'Nguyễn Danh Lâm và cộng sự (2022) sử dụng thang đo DASS-21 khảo sát trên 482 học sinh trung học phổ thông (2 trường) tại huyện Yên Định, Thanh Hóa — huyện bán nông thôn — báo cáo tỷ lệ lo âu chung 49,0% (nhẹ 11,2%; vừa 25,1%; nặng 8,1%; rất nặng 4,6%). Nghiên cứu này không phân tích lo âu theo giới tính.'),
    # An Giang
    ('Lê Minh T. và cộng sự (2025 ) khảo sát 600 học sinh phổ thông Long Bình, An Giang bằng DASS-21 báo cáo tỷ lệ 33,2%.',
     'Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh (2025) khảo sát 366 học sinh THCS&THPT Long Bình, An Giang bằng DASS-21 báo cáo tỷ lệ lo âu 61,2% (nhẹ 9,3%; vừa 24,0%; nặng 12,6%; rất nặng 15,3%).'),
    # Geographic update: An Giang context
    ('Yên Định (bán đô thị)', 'Yên Định (bán nông thôn)'),
]


def apply_repl(p, replacements, applied):
    """Apply replacement in paragraph, marked red."""
    full = p.text
    if not full: return False
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            i = full.find(old, cursor)
            if i == -1: break
            matches.append((i, i + len(old), old, new))
            cursor = i + len(old)
    if not matches: return False
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    safe = []
    last = -1
    for m in matches:
        if m[0] >= last:
            safe.append(m); last = m[1]
    segments = []
    cur = 0
    for s, e, o, n in safe:
        if cur < s:
            segments.append((full[cur:s], False))
        segments.append((n, True))
        applied.append(f"'{o[:60]}...' -> '{n[:60]}...'")
        cur = e
    if cur < len(full):
        segments.append((full[cur:], False))
    for r in p.runs:
        r.text = ''
    for text, is_red in segments:
        if not text: continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


d = Document(FILE)
applied = []
for i, p in enumerate(d.paragraphs):
    if apply_repl(p, REPLACEMENTS, applied):
        print(f"Fixed para {i}:")
        for a in applied[-5:]:
            print(f"   ✓ {a}")

d.save(FILE)
print(f"\nSaved: {FILE}")
print(f"Total fixes: {len(applied)}")
