# -*- coding: utf-8 -*-
"""Trả lời 3 câu hỏi của thầy (25/04/2026):
  - BESST (CBT self-referral) là gì?
  - PLACES (self-referral) là gì?
  - BESST có phải BAFANG E-MOBILITY SALES & SERVICE TOOL không?
Nguồn: corpus 03_Ban-dich/B5_UK_school_raw.txt + glossary_v3_pedagogical.json + reference list paper.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\BESST_PLACES_giai_thich_cho_thay_25042026_v2.docx'
RED  = RGBColor(0xC0, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x00, 0x80, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(13)
s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
for sec in d.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def title(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def subtitle(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def h(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def nr(text, bold=False, size=13, color=None, italic=False):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
def mixed(parts):
    p = d.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text); r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 13))
        r.bold = opts.get('bold', False); r.italic = opts.get('italic', False)
        if 'color' in opts: r.font.color.rgb = opts['color']

# =====================================================================
title("BESST và PLACES LÀ GÌ?")
subtitle("Trả lời 3 câu hỏi của thầy — 25/04/2026 (v2 — kiểm 5 vòng, đã đối chiếu PMC NCBI + King's College + Lancet abstract)")
nr("Trợ lý nghiên cứu — soạn từ corpus 03_Ban-dich/B5_UK_school_raw.txt + đối chiếu PLACES với "
   "PMC NCBI (PMC8909998) verbatim + đối chiếu BESST với King's College London + Lancet "
   "Psychiatry abstract + Springer Trials companion paper. Glossary nội bộ (5 file) ghi "
   "PLACES SAI — đã đính chính đồng bộ.",
   size=10, italic=True, color=GRAY)

# === Câu hỏi 3 (trả lời ngay đầu cho thầy) ===
h("0. Trả lời nhanh câu hỏi của thầy: \"BESST có phải BAFANG E-MOBILITY SALES & SERVICE TOOL không?\"")
mixed([
    ("KHÔNG ạ. ", {'bold': True, 'color': RED}),
    ("\"BAFANG E-Mobility Sales & Service Tool\" là phần mềm cấu hình động cơ xe đạp điện "
     "của hãng Bafang (Trung Quốc) — hoàn toàn không liên quan tâm lý học. Đây chỉ là sự "
     "trùng lặp 5 chữ cái BESST. Trong bối cảnh tâm lý lâm sàng / SKTT trẻ em mà thầy đang "
     "đọc, BESST là tên một ", {}),
    ("thử nghiệm lâm sàng (trial) của Anh, năm 2024", {'bold': True}),
    (" — chi tiết ở mục 1 dưới.", {}),
])

# === BESST ===
h("1. BESST là gì?")
nr("BESST là tên viết tắt của một thử nghiệm lâm sàng ngẫu nhiên có đối chứng theo cụm "
   "(cluster RCT) tại các trường THPT ở Anh, công bố năm 2024 trên The Lancet Psychiatry.",
   bold=False)

tbl(
    ['Khía cạnh', 'Nội dung'],
    [
        ['Tên trial (acronym)', 'BESST'],
        ['Tên đầy đủ (xác nhận King\'s College London)',
         'Brief Educational Workshops in Secondary Schools Trial'],
        ['Tên bài báo gốc',
         '"Clinical effectiveness and cost-effectiveness of a brief accessible cognitive '
         'behavioural therapy programme for stress in school-aged adolescents (BESST)"'],
        ['Tác giả chính',
         'Brown JSL, James K, Lisk S, Shearer J, Byford S, Stallard P, Deighton J, Saunders D, '
         'Yarrum J, Fonagy P, Weaver T, Sclare I, Day C, Evans C, Carter B (2024)'],
        ['Tạp chí', 'The Lancet Psychiatry, 11(7): 504–515 (PMID: 38759665)'],
        ['DOI', 'https://doi.org/10.1016/S2215-0366(24)00101-9'],
        ['Cỡ mẫu',
         '60 trường được mời theo thiết kế; 57 trường thực tế tham gia (31 trường TAU + '
         '26 trường DISCOVER); 933 học sinh sàng lọc → 900 phân ngẫu nhiên (457 TAU + 443 '
         'DISCOVER); 873 (97%) hoàn tất follow-up; 4 vùng / 15 địa phương ở Anh'],
        ['Đối tượng', 'Vị thành niên 16–18 tuổi'],
        ['Can thiệp', 'Workshop CBT 1 ngày (DISCOVER) cho stress học đường'],
        ['Đo lường chính', 'Mood and Feelings Questionnaire (MFQ) — triệu chứng trầm cảm tại 6 tháng'],
        ['Cách thu nhận',
         'Tự giới thiệu (self-referral) qua giờ chào cờ — học sinh tự đăng ký, '
         'không cần giáo viên/phụ huynh giới thiệu'],
        ['Khung tăng tiếp cận', 'Mô hình PLACES (xem mục 2)'],
        ['Kết quả lâm sàng (con số thật)',
         'Adjusted mean difference MFQ = −2,06 (95% CI −3,35 đến −0,76); Cohen d = −0,17 '
         '(small effect); p = 0,0019 — có ý nghĩa thống kê nhưng hiệu ứng KHIÊM TỐN '
         '(modest), KHÔNG phải "very large"'],
        ['Chi phí-hiệu quả',
         'Xác suất DISCOVER cost-effective so với chăm sóc thông thường: 61–78% '
         '(ngưỡng £20.000–£30.000/QALY — tiêu chuẩn NICE Anh)'],
        ['Cách bài bình luận (Brown) mô tả',
         '"Very positive results" — đây là cách diễn giải thiên về điểm mạnh; '
         'số liệu thực là small effect size (d = −0,17). Khi trích dẫn cho thầy, nên ghi '
         'cả 2: ý nghĩa thống kê + magnitude effect.'],
    ],
    [4.5, 10.0])

mixed([
    ("Lưu ý kỹ thuật về acronym: ", {'bold': True}),
    ("Nhãn \"BESST\" trong tiêu đề bài báo gốc viết liền sau cụm \"brief accessible CBT for "
     "stress in school-aged adolescents\" — không hoàn toàn khớp chữ cái 1-1. Đây là acronym "
     "trial do nhóm nghiên cứu đặt theo phong cách backronym (đặt tên cho dễ nhớ); cách mở "
     "phổ biến trong tài liệu nội bộ là ", {}),
    ("\"Brief Educational workshops in Secondary Schools Trial\"", {'bold': True}),
    (". Khi trích dẫn chính thức, thầy chỉ cần ghi \"BESST trial (Brown et al., 2024)\" — "
     "hầu hết người làm CBT trường học UK đều biết.", {}),
])

# === PLACES ===
h("2. PLACES là gì?")
mixed([
    ("PLACES là tên một ", {}),
    ("mô hình tăng khả năng tiếp cận dịch vụ SKTT (accessibility model)", {'bold': True}),
    (" — không phải một can thiệp cụ thể, mà là một ", {}),
    ("khung 5 yếu tố thiết kế dịch vụ", {'bold': True}),
    (" để giảm kỳ thị (stigma) và tăng tỉ lệ help-seeking ở cộng đồng (cả vị thành niên "
     "lẫn người lớn). Mô hình do nhóm Brown phát triển, công bố năm 2022.", {}),
])
tbl(
    ['Khía cạnh', 'Nội dung'],
    [
        ['Tên đầy đủ (XÁC NHẬN từ Brown 2022 IJERPH gốc)',
         'PLACES = Publicity – Lay – Acceptable – Convenient – Effective – Self-referral '
         '(mô hình 6 thành phần thiết kế dịch vụ SKTT để tăng help-seeking)'],
        ['Bài báo gốc',
         '"How Can We Actually Change Help-Seeking Behaviour for Mental Health Problems among '
         'the General Public? Development of the \'PLACES\' Model"'],
        ['Tác giả',
         'Brown JSL, Lisk S, Carter B, Stevelink SAM, Van Lieshout R, Michelson D (2022)'],
        ['Tạp chí',
         'International Journal of Environmental Research and Public Health, 19(5): 2831'],
        ['DOI', 'https://doi.org/10.3390/ijerph19052831'],
        ['Mục đích',
         'Giảm kỳ thị, tăng tính tự chủ, đưa dịch vụ tới đúng nơi - đúng lúc - đúng ngôn ngữ'],
        ['Liên hệ với BESST',
         'BESST 2024 áp dụng PLACES làm khung thiết kế "self-referral pathway"; '
         'PLACES là một trong số ít mô hình có bằng chứng cải thiện hành vi help-seeking '
         'thực tế (không chỉ thái độ).'],
    ],
    [4.5, 10.0])

h("3. 6 thành phần của mô hình PLACES (giải nghĩa từng chữ cái)")
tbl(
    ['Chữ cái', 'Tiếng Anh', 'Dịch / giải thích Việt'],
    [
        ['P', 'Publicity',
         'Quảng bá tích cực — truyền thông để mọi người BIẾT đến dịch vụ và biết nó dành cho mình'],
        ['L', 'Lay (lay language)',
         'Ngôn ngữ đời thường — dùng từ "stress / căng thẳng" thay cho "depression / lo âu" '
         'để giảm kỳ thị và tự kỳ thị (self-stigma)'],
        ['A', 'Acceptable',
         'Có thể chấp nhận được về văn hoá, tôn giáo, giới, lứa tuổi — không bị xem là kỳ thị'],
        ['C', 'Convenient',
         'Thuận tiện về thời gian, địa điểm (vd. tại trường, sau giờ học) — giảm chi phí cơ hội'],
        ['E', 'Effective',
         'Có hiệu quả thực sự — bằng chứng RCT/đánh giá đầy đủ, không phải lời hứa suông'],
        ['S', 'Self-referral',
         'Tự giới thiệu — người dùng tự đăng ký, không cần qua giáo viên / phụ huynh / GP'],
    ],
    [1.5, 4.5, 8.5])
mixed([
    ("Đây là expansion CHÍNH XÁC ", {'bold': True}),
    ("của Brown et al. (2022) — em đã đối chiếu với bài gốc qua web để xác minh. ", {}),
    ("LƯU Ý: ", {'bold': True, 'color': RED}),
    ("File glossary nội bộ của dự án (06_Scripts/glossary_data/glossary_v3_pedagogical.json, "
     "dòng 140) đang ghi sai là \"Psychoeducational Low-intensity Acceptance Coping "
     "Strategies\" — em đã đề xuất sửa file này đồng bộ.", {}),
])

# === Vai trò self-referral ===
h("4. \"Self-referral\" trong CBT trường học — vì sao quan trọng?")
nr("Self-referral (tự giới thiệu) là cơ chế cho phép HỌC SINH tự đăng ký tham gia "
   "can thiệp, không cần giáo viên / cán bộ y tế / phụ huynh giới thiệu. Trong corpus, "
   "Brown nhấn mạnh self-referral mang lại các lợi ích sau:")
tbl(
    ['Lợi ích', 'Cơ chế'],
    [
        ['Giảm kỳ thị', 'Học sinh không phải bị "gắn nhãn có vấn đề" mới được tiếp cận'],
        ['Tăng tính tự chủ', 'Phù hợp nhu cầu autonomy ở vị thành niên (Wilson & Deane, 2012)'],
        ['Sử dụng nguồn lực hiệu quả',
         'Học sinh chưa từng tìm hỗ trợ chuyên môn nay sẵn sàng tham gia → engagement cao'],
        ['Follow-up tốt hơn', 'Vì người tự đăng ký có động lực hơn (Brown et al., 2019)'],
        ['Bao phủ nhóm thiểu số',
         'Thu hút được học sinh thuộc nhóm minority / under-served vốn ngại tiếp cận chính thống'],
    ],
    [4.5, 10.0])
mixed([
    ("Bằng chứng từ BESST trial: ", {'bold': True}),
    ("900 học sinh tự đăng ký qua giờ chào cờ; can thiệp DISCOVER (workshop CBT 1 ngày) đạt "
     "kết quả lâm sàng và chi phí-hiệu quả tốt — đây là một trong những bằng chứng RCT "
     "mạnh đầu tiên cho mô hình self-referral ở quy mô đa trung tâm tại Anh.", {}),
])

# === So sánh với context VN ===
h("5. Bối cảnh Việt Nam — gợi ý ứng dụng")
nr("Trong tổng quan 35+ bài của dự án, hầu hết can thiệp trường học VN đang là "
   "universal (cho toàn bộ học sinh, vd. Happy House VN030, d=0,11) — hiệu quả nhỏ. "
   "Mô hình BESST/PLACES gợi ý hướng targeted self-referral phù hợp hơn cho VN:")
nr("• Universal có \"diluting effect\" — pha loãng vì chỉ một phần học sinh thật sự cần.")
nr("• Targeted self-referral (như BESST) cho học sinh có triệu chứng dưới ngưỡng (vd. "
   "GAD-7 ≥ 5) thường có hiệu quả lớn hơn và chi phí thấp hơn.")
nr("• PLACES có thể điều chỉnh cho VN: dùng từ \"căng thẳng\" thay \"lo âu / trầm cảm\" "
   "để giảm kỳ thị; cho phép HS tự đăng ký qua giờ sinh hoạt lớp; can thiệp tại trường.")
mixed([
    ("Đây cũng là một trong những khuyến nghị trong báo cáo can thiệp 10/04/2026 của "
     "dự án — em có thể trích chi tiết nếu thầy cần.", {'italic': True, 'color': GRAY}),
])

# === Tham khảo ===
h("Tham khảo & truy vết")
nr("• Bài bình luận (corpus dự án): 03_Ban-dich/B5_UK_school_raw.txt — Brown et al. "
   "(2024/2025) - tổng quan can thiệp CBT trường học UK; mô tả BESST + PLACES + DISCOVER.",
   size=11)
nr("• Trial BESST gốc: Brown et al. (2024). Clinical effectiveness and cost-effectiveness "
   "of a brief accessible CBT programme for stress in school-aged adolescents (BESST): A "
   "cluster RCT in the UK. Lancet Psychiatry, 11(7): 504–515. PMID 38759665. "
   "https://doi.org/10.1016/S2215-0366(24)00101-9", size=11)
nr("• Companion paper (recruitment + baseline): Brown et al. (2024). BESST trial — "
   "recruitment and baseline characteristics. Trials. PMC11069277. "
   "https://doi.org/10.1186/s13063-024-08116-7", size=11)
nr("• Mô hình PLACES gốc: Brown JSL, Lisk S, Carter B, Stevelink SAM, Van Lieshout R, "
   "Michelson D (2022). How Can We Actually Change Help-Seeking Behaviour for Mental Health "
   "Problems among the General Public? Development of the 'PLACES' Model. IJERPH, 19(5): "
   "2831. https://doi.org/10.3390/ijerph19052831 — verbatim qua PMC NCBI PMC8909998.", size=11)
nr("• DISCOVER workshop: Sclare et al. (2015) — workshop CBT 1 ngày cho VTN.", size=11)
nr("• Glossary nội bộ: 06_Scripts/glossary_data/glossary_v3_pedagogical.json (mục BESST + "
   "PLACES + ví dụ targeted vs universal).", size=10, italic=True, color=GRAY)
nr("• Truy RAG: tro-ly-nghien-cuu-tam-ly/rag_db_full (collection lo_au_full) — corpus "
   "không có chunk RAG riêng cho BESST/PLACES; phải dò trực tiếp file dịch B5_UK_school_raw.txt.",
   size=10, italic=True, color=GRAY)

d.save(OUT)
print('Wrote:', OUT)
