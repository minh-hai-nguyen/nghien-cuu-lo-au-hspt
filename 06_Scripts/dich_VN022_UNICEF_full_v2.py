# -*- coding: utf-8 -*-
"""
VN022 UNICEF 2022 — Dich FULL tung cau tung chu theo template Jenkins
Input: PDF 118 trang goc
Output: 03_Ban-dich/VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx

BATCH 1: Front matter (p1-5) + Chapter 1 Introduction (p6-10)
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DST = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(12)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.5
for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)

PAGE_W = 16.0
RED = RGBColor(0xCC, 0, 0)

# ============ helpers ============
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def H(text, level=1, red=False):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RED if red else RGBColor(0, 0, 0)

def Pred(text, bold=False, italic=False, size=12, align='justify'):
    return P(text, bold=bold, italic=italic, size=size, color=RED, align=align)

def P(text, bold=False, italic=False, size=12, color=None, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'left':
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color: r.font.color.rgb = color

def PageMark(text):
    """--- Trang X–Y, UNICEF Viet Nam, 2022 ---"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/unicef_imgs/'
def AddImg(filename, caption_vn, caption_en, width_cm=11.0):
    """Add image with bilingual caption (Nguyen tac 3 — bao toan anh)."""
    path = IMG_DIR + filename
    if not os.path.exists(path):
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try:
        p.add_run().add_picture(path, width=Cm(width_cm))
    except Exception as e:
        return
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def table2col(pairs, widths=None):
    if widths is None: widths = [5.0, 11.0]
    t = doc.add_table(rows=len(pairs), cols=2)
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_grid(t, widths)
    for ri, (k, v) in enumerate(pairs):
        for ci, val in enumerate([k, v]):
            c = t.rows[ri].cells[ci]; c.text = val
            colw(c, widths[ci])
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
                    if ci == 0:
                        r.bold = True
                        shade(c, 'D9E2F3')

# ============================================================
# BIA + TIEU DE
# ============================================================
P('Link bản gốc: https://www.unicef.org/vietnam/reports/study-school-related-factors-'
  'impacting-mental-health-and-well-being-adolescents-viet-nam',
  align='center', italic=True, size=10)
P('Quét mã QR để truy cập báo cáo gốc', align='center', italic=True, size=10)

H('NGHIÊN CỨU TOÀN DIỆN VỀ CÁC YẾU TỐ LIÊN QUAN ĐẾN TRƯỜNG HỌC ẢNH HƯỞNG ĐẾN SỨC KHOẺ '
  'TÂM THẦN VÀ HẠNH PHÚC CỦA THANH THIẾU NIÊN NAM VÀ NỮ TẠI VIỆT NAM', level=1)
P('(Comprehensive Study on School-Related Factors Impacting Mental Health and Well-Being '
  'of Adolescent Boys and Girls in Viet Nam)', italic=True, align='center')
P('UNICEF Việt Nam, 2022 — Báo cáo chính sách, 118 trang', align='center', italic=True)

# ============================================================
# --- Trang 1, UNICEF Viet Nam, 2022 ---
# ============================================================
PageMark('--- Trang 1, UNICEF Việt Nam, 2022 ---')
P('Tiêu đề đầy đủ (bìa báo cáo):', bold=True)
P('NGHIÊN CỨU TOÀN DIỆN VỀ CÁC YẾU TỐ LIÊN QUAN ĐẾN TRƯỜNG HỌC ẢNH HƯỞNG ĐẾN SỨC KHOẺ '
  'TÂM THẦN VÀ HẠNH PHÚC CỦA THANH THIẾU NIÊN NAM VÀ NỮ TẠI VIỆT NAM',
  align='center', bold=True)
AddImg('p001_img1_Im0.jpg',
       'Hình 1 (bìa). Học sinh dân tộc thiểu số Việt Nam',
       'Cover image — Vietnamese ethnic minority adolescent', width_cm=12.0)

# ============================================================
# --- Trang 2-3, UNICEF Viet Nam, 2022 ---
# ============================================================
PageMark('--- Trang 2–3, UNICEF Việt Nam, 2022 ---')

H('LỜI CẢM ƠN (Acknowledgements)', level=2)

P('Nghiên cứu mang tên "Nghiên cứu toàn diện về các yếu tố liên quan đến trường học ảnh '
  'hưởng đến sức khoẻ tâm thần và hạnh phúc của thanh thiếu niên nam và nữ tại Việt Nam" '
  'được thực hiện như một phần của hợp tác rộng hơn giữa UNICEF Việt Nam và Bộ Giáo dục '
  'và Đào tạo (MOET), với chuyên môn nghiên cứu và kỹ thuật do Tư vấn Quốc tế Amie Alley '
  'Pollack, Tư vấn Quốc gia Đặng Hoàng Minh (Đại học Quốc gia Việt Nam) và các Trợ lý '
  'Nghiên cứu là Phương Nguyễn, Trang Lê và Lâm Lê cung cấp. Nghiên cứu nhằm xây dựng hiểu '
  'biết về việc các yếu tố liên quan đến trường học tác động như thế nào đến sức khoẻ tâm '
  'thần và hạnh phúc của thanh thiếu niên nam và nữ tại Việt Nam, cũng như vai trò của hệ '
  'thống giáo dục trong việc giải quyết cả các rủi ro liên quan đến trường học và việc '
  'cung cấp dịch vụ hỗ trợ sức khoẻ tâm thần và tâm lý xã hội.')

P('Nghiên cứu đã nhận được hỗ trợ tài chính và phương pháp từ UNICEF Việt Nam. Nhóm '
  'nghiên cứu xin ghi nhận đóng góp to lớn từ đội ngũ Ban Giáo dục của UNICEF Việt Nam '
  'vì sự hợp tác không mệt mỏi, hỗ trợ, tổ chức các hoạt động nghiên cứu và rà soát tất '
  'cả các bản thảo báo cáo. Nhóm nghiên cứu cũng xin ghi nhận sự hỗ trợ của các chuyên '
  'gia phản biện bên ngoài: Lesley Miller, Ngô Thị Quỳnh Hoa, Marcy Levy, Roshni Basu, '
  'Joanna Lai, Lê Anh Lan, Rachel Harvey và Vincenzo Vinci, vì những nhận xét chu đáo '
  'và chi tiết trên các bản thảo đầu. Nhóm nghiên cứu đánh giá cao các Vụ Giáo dục Chính '
  'trị và Công tác Học sinh – Sinh viên và Vụ Giáo dục Thể chất của Bộ GD&ĐT đã hỗ trợ '
  'nghiên cứu, bao gồm tham vấn quan trọng trên các bản thảo báo cáo. Và cuối cùng, nhóm '
  'nghiên cứu xin cảm ơn rất nhiều học sinh, phụ huynh, giáo viên, quản lý và chuyên gia '
  'đã tham gia đóng góp thời gian và trả lời nhiều câu hỏi.')

# ============================================================
# --- Trang 4, UNICEF Viet Nam, 2022 ---
# ============================================================
PageMark('--- Trang 4, UNICEF Việt Nam, 2022 ---')

H('BẢNG VIẾT TẮT (Abbreviations)', level=2)

table2col([
    ('ADHD', 'Attention Deficit Hyperactivity Disorder — Rối loạn Tăng động Giảm chú ý'),
    ('CAMH', 'Child and Adolescent Mental Health — Sức khoẻ Tâm thần Trẻ em và Vị thành niên'),
    ('CBT', 'Cognitive Behaviour Therapy — Liệu pháp Nhận thức Hành vi'),
    ('DOET', 'Department of Education and Training — Sở Giáo dục và Đào tạo'),
    ('DOH', 'Department of Health — Sở Y tế'),
    ('DOLISA', 'Department of Labour, Invalids and Social Affairs — Sở Lao động, Thương binh và Xã hội'),
    ('FGD', 'Focus Group Discussion — Thảo luận Nhóm tập trung'),
    ('HCMC', 'Ho Chi Minh City — Thành phố Hồ Chí Minh'),
    ('KII', 'Key Informant Interview — Phỏng vấn Người cung cấp thông tin chủ chốt'),
    ('LGBTQ', 'Lesbian, Gay, Bisexual, Transgender and Queer (or Questioning) — Đồng tính nữ, đồng tính nam, song tính, chuyển giới và queer (hoặc đang tìm hiểu bản dạng)'),
    ('MHPSS', 'Mental Health and Psychosocial Support — Hỗ trợ Sức khoẻ Tâm thần và Tâm lý xã hội'),
    ('MOET', 'Ministry of Education and Training — Bộ Giáo dục và Đào tạo'),
    ('MOH', 'Ministry of Health — Bộ Y tế'),
    ('MOLISA', 'Ministry of Labour, Invalids and Social Affairs — Bộ Lao động, Thương binh và Xã hội'),
    ('PED', 'Physical Education Department, Ministry of Education and Training — Vụ Giáo dục Thể chất, Bộ GD&ĐT'),
    ('PESAD', 'Political Education and Student Affairs Department, Ministry of Education and Training — Vụ Giáo dục Chính trị và Công tác Học sinh – Sinh viên, Bộ GD&ĐT'),
    ('PTSD', 'Post-Traumatic Stress Disorder — Rối loạn Stress Sau sang chấn'),
    ('UNICEF', 'United Nations Children\'s Fund — Quỹ Nhi đồng Liên Hợp Quốc'),
])

# ============================================================
# --- Trang 5, UNICEF Viet Nam, 2022 ---
# ============================================================
PageMark('--- Trang 5, UNICEF Việt Nam, 2022 ---')

H('MỤC LỤC (Table of Content)', level=2)

toc_items = [
    ('Lời cảm ơn (Acknowledgements)', '1'),
    ('Bảng viết tắt (Abbreviations)', '2'),
    ('Chương 1: Giới thiệu (Introduction)', '4'),
    ('Chương 2: Sức khoẻ tâm thần vị thành niên — Tổng quan tài liệu toàn cầu '
     '(Adolescent Mental Health: Global Literature Review)', '9'),
    ('Chương 3: Tổng quan tài liệu về sức khoẻ tâm thần vị thành niên tại Việt Nam '
     '(Adolescent Mental Health in Viet Nam Literature Review)', '21'),
    ('Chương 4: Phương pháp (Methodology)', '26'),
    ('Chương 5: Phát hiện nghiên cứu về sức khoẻ tâm thần và hạnh phúc của học sinh '
     'vị thành niên Việt Nam (Adolescent Student Mental Health and Well-Being in '
     'Viet Nam Study Findings)', '33'),
    ('Chương 6: Các yếu tố liên quan đến trường học ảnh hưởng đến sức khoẻ tâm thần '
     'và hạnh phúc học sinh tại Việt Nam (School-Related Factors Impacting Student '
     'Mental Health and Well-being in Viet Nam)', '41'),
    ('Chương 7: Chính sách và Chương trình MHPSS học đường cho học sinh vị thành niên '
     'tại Việt Nam (School-Based Adolescent Student MHPSS Policies and Programmes '
     'in Viet Nam)', '53'),
    ('Chương 8: Chính sách và Chương trình MHPSS cho học sinh vị thành niên tại Việt '
     'Nam — Khu vực xã hội và y tế, hợp tác liên ngành (Adolescent Student MHPSS '
     'Policies and Programmes in Viet Nam: Social and Health Sectors, and Cross-'
     'Sector Collaboration)', '64'),
    ('Chương 9: Kết luận và Khuyến nghị (Conclusions and Recommendations)', '73'),
    ('Phụ lục 1 (Appendix 1)', '82'),
    ('Phụ lục 2 (Appendix 2)', '86'),
    ('Phụ lục 3 (Appendix 3)', '90'),
    ('Tài liệu tham khảo (References)', '108'),
]
table2col([(f'{k}', f'Trang {v}') for k, v in toc_items], widths=[13.0, 3.0])

P('** Vui lòng tham khảo Tóm tắt điều hành (Executive Summary) để có bản tóm tắt ngắn '
  'gọn về các phát hiện và khuyến nghị của báo cáo.', italic=True, size=11)
P('(Ghi chú của người dịch: Tóm tắt điều hành là tài liệu đi kèm riêng, không nằm '
  'trong 118 trang báo cáo chính này.)', italic=True, size=10, color=RGBColor(0x66, 0x66, 0x66))

doc.add_page_break()

# ============================================================
# CHUONG 1 — GIOI THIEU
# ============================================================
H('CHƯƠNG 1: GIỚI THIỆU (Introduction)', level=1)

# --- Trang 6, UNICEF, 2022 — Study Problem and Rationale ---
PageMark('--- Trang 6, UNICEF Việt Nam, 2022 ---')

H('Vấn đề và lý do nghiên cứu (Study Problem and Rationale)', level=2)

P('Trên toàn thế giới, khoảng 15 % trẻ em và vị thành niên mắc các rối loạn tâm thần, '
  'và các bệnh lý tâm thần kinh là nguyên nhân hàng đầu gây tàn tật ở người trẻ '
  '(Polanczyk, Salum, Sugaya, Caye & Rohde, 2015). Với 50 % rối loạn sức khoẻ tâm '
  'thần khởi phát trước tuổi 14 và 75 % trước tuổi 24 (Kessler và cộng sự, 2005), sức '
  'khoẻ tâm thần trẻ em và vị thành niên đã trở thành ưu tiên toàn cầu. Có một khoảng '
  'cách đáng kể về giới trên toàn thế giới trong sức khoẻ tâm thần vị thành niên, với '
  'các em gái trung bình có sức khoẻ tâm thần kém hơn các em trai (Campbell, Bann & '
  'Patalay, 2021). Tại Việt Nam, bằng chứng cho thấy 8 % – 29 % vị thành niên gặp các '
  'vấn đề sức khoẻ tâm thần, trong đó các em trai có tỷ lệ rối loạn hành vi cao hơn và '
  'các em gái có tỷ lệ các vấn đề cảm xúc cao hơn như lo âu và trầm cảm (UNICEF, 2018; '
  'Weiss và cộng sự, 2014). Các vấn đề sức khoẻ tâm thần là gánh nặng đáng kể đối với '
  'vị thành niên tại Việt Nam. Ở các em trai, rối loạn hành vi là nguyên nhân gây tàn '
  'tật đứng thứ hai trong nhóm 10–14 tuổi và thứ năm trong nhóm nam 15–19 tuổi. Các em '
  'gái đối mặt với nhiều nguồn gây tàn tật liên quan đến sức khoẻ tâm thần, với cả rối '
  'loạn hành vi và rối loạn lo âu đều là nguyên nhân gây tàn tật hàng đầu trong nhóm '
  '10–14 tuổi, và rối loạn trầm cảm là nguyên nhân gây tàn tật đứng thứ ba trong nhóm '
  'nam 15–19 tuổi (UNICEF Country Dashboard, 2019). Sự hiểu biết kém về các vấn đề sức '
  'khoẻ tâm thần, kỳ thị xã hội và dịch vụ cũng như nguồn lực sức khoẻ tâm thần hạn chế '
  'khiến phần lớn trẻ em này không nhận được điều trị hay hỗ trợ. Khi các vấn đề sức '
  'khoẻ tâm thần mới chớm nở bị bỏ không điều trị, những vấn đề này tác động nghiêm '
  'trọng đến sự phát triển, thành tích học tập và tiềm năng trong cuộc sống của trẻ em.')

P('Vị thành niên (10–19 tuổi) là giai đoạn quan trọng để phát triển các kỹ năng cảm '
  'xúc – xã hội thiết yếu cho chức năng lành mạnh và hạnh phúc. Gia đình, trường học '
  'và cộng đồng hỗ trợ giúp vị thành niên có thể ứng phó hiệu quả với những thách thức '
  'cuộc sống và phát triển các kỹ năng tâm lý – xã hội lành mạnh. Các yếu tố liên quan '
  'đến trường học vừa có thể làm tăng nguy cơ mắc vấn đề sức khoẻ tâm thần, vừa có '
  'thể bảo vệ sức khoẻ và hạnh phúc của vị thành niên. Bằng chứng rộng rãi chỉ ra rằng '
  'khí hậu trường học (school climate), bao gồm các yếu tố an toàn, gắn kết và môi '
  'trường, tác động đến sức khoẻ tâm thần vị thành niên (McChesney & Aldridge, 2018). '
  'Áp lực học tập — từ phía cha mẹ, giáo viên và bạn bè, và liên quan đến chương trình '
  'học dày đặc cùng yêu cầu thi cử — là một yếu tố khác liên quan đến trường học gắn '
  'với các vấn đề sức khoẻ tâm thần vị thành niên (Nguyen, Dedding, Pham, Wright & '
  'Bunders, 2013). Bắt nạt và các nguồn gây stress xã hội khác đã được chứng minh có '
  'tác động tiêu cực đến sức khoẻ tâm thần ở học sinh vị thành niên tại Việt Nam (Le '
  'và cộng sự, 2019).')

P('Các dịch vụ sức khoẻ tâm thần học đường (school-based mental health services) '
  'đóng vai trò quan trọng trong việc hỗ trợ sức khoẻ tâm thần học sinh vị thành niên '
  'và giải quyết các yếu tố nguy cơ sức khoẻ tâm thần liên quan đến trường học. Trường '
  'học là một trong những môi trường tâm lý – xã hội quan trọng nhất của vị thành niên, '
  'cho phép tiếp cận trực tiếp với các nhân vật trung tâm trong đời sống các em (giáo '
  'viên, bạn bè). Hỗ trợ sức khoẻ tâm thần và tâm lý xã hội học đường (School-Based '
  'MHPSS) cho phép tiếp cận tương đối dễ dàng các học sinh cần được chăm sóc, vì dịch '
  'vụ không phụ thuộc vào việc cha mẹ hay người giám hộ phải đưa trẻ đến phòng khám. '
  'Và khi triển khai đúng cách, các dịch vụ này có thể giảm kỳ thị đối với dịch vụ '
  'sức khoẻ tâm thần bằng cách tích hợp chúng vào giáo dục. Nghiên cứu về việc sử '
  'dụng các can thiệp sức khoẻ tâm thần học đường phổ quát (tức là nhắm vào tất cả học '
  'sinh) và chọn lọc (tức là nhắm vào học sinh có nguy cơ) cho thấy trường học cung '
  'cấp môi trường tuyệt vời để nhắm vào sức khoẻ tâm thần của trẻ, thành tích học tập '
  'và mối liên kết quan trọng giữa hai yếu tố này (Greenwood, Kratochwill & Clements, '
  '2008). Tại Việt Nam, nghiên cứu ủng hộ việc sử dụng các chương trình học đường để '
  'hỗ trợ sức khoẻ tâm thần trẻ em (Dang, Weiss, Nguyen, Tran & Pollack, 2017).')

# --- Trang 7, UNICEF, 2022 — Significance of the Study ---
PageMark('--- Trang 7, UNICEF Việt Nam, 2022 ---')

H('Ý nghĩa của nghiên cứu (Significance of the Study)', level=2)

P('Nghiên cứu này phân tích dữ liệu gốc (primary data) về sức khoẻ tâm thần học sinh '
  'vị thành niên, bao gồm nhu cầu sức khoẻ tâm thần và các yếu tố nguy cơ liên quan đến '
  'trường học. Nghiên cứu cũng bao gồm rà soát các chính sách và chương trình hiện có '
  'liên quan đến sức khoẻ tâm thần học sinh vị thành niên. Phân tích dữ liệu này đem '
  'lại những hiểu biết về nhu cầu sức khoẻ tâm thần học sinh và các lĩnh vực tiềm năng '
  'cho hỗ trợ về mặt chương trình và chính sách. Cụ thể, các phát hiện và khuyến nghị '
  'của nghiên cứu sẽ đặc biệt hữu ích cho UNICEF Việt Nam và Bộ Giáo dục và Đào tạo '
  '(MOET) của Việt Nam, đồng thời có giá trị thứ cấp cho Bộ Y tế (MOH) và Bộ Lao động, '
  'Thương binh và Xã hội (MOLISA).')

P('UNICEF đã đưa sức khoẻ tâm thần và hạnh phúc của trẻ em và vị thành niên trở thành '
  'một ưu tiên chiến lược then chốt cho giai đoạn 2022–2026. Trọng tâm này phản ánh '
  'vai trò thiết yếu của sức khoẻ tâm thần trong sự phát triển của trẻ, sự gắn kết và '
  'thành công trong học tập, các mối quan hệ lành mạnh, và sự tham gia đầy đủ vào cuộc '
  'sống và xã hội. Với đại dịch toàn cầu vẫn đang đặt ra những thách thức to lớn đối '
  'với trẻ em và thanh thiếu niên, việc ưu tiên các nhu cầu sức khoẻ tâm thần trẻ em '
  'trở nên cần thiết hơn bao giờ hết. Nghiên cứu này sẽ cung cấp thông tin cho việc lập '
  'kế hoạch chiến lược và phát triển chương trình của UNICEF Việt Nam khi họ làm việc '
  'để hỗ trợ chính phủ Việt Nam trong lĩnh vực này.')

P('Bộ GD&ĐT đã thể hiện cam kết lớn trong việc giải quyết nhu cầu sức khoẻ tâm thần '
  'của học sinh vị thành niên. Các phát hiện và khuyến nghị của nghiên cứu sẽ hỗ trợ '
  'Bộ GD&ĐT khi họ xem xét chính sách hỗ trợ hạnh phúc học sinh, và khi họ làm việc '
  'để phát triển các chương trình và dịch vụ hiệu quả cho học sinh trong trường học. '
  'Và cuối cùng, các phát hiện của nghiên cứu sẽ cung cấp thông tin cho tất cả các bộ '
  'ngành tham gia hỗ trợ sức khoẻ tâm thần vị thành niên — MOET, MOH và MOLISA — khi '
  'họ làm việc để xem xét các lĩnh vực có khả năng hợp tác và điều phối cung cấp chính '
  'sách và dịch vụ. Mỗi ngành đóng vai trò thiết yếu đối với sức khoẻ, sự phát triển và '
  'thành công tương lai của học sinh vị thành niên. Do đó, chúng tôi hy vọng nghiên '
  'cứu sẽ cung cấp thông tin cho mỗi Bộ và hỗ trợ hợp tác cấp Bộ ngành.')

# --- Trang 8, UNICEF, 2022 — Key Audiences ---
PageMark('--- Trang 8, UNICEF Việt Nam, 2022 ---')
AddImg('p008_img1_Im0.jpg',
       'Hình 2 (trang 8). Học sinh nữ vị thành niên — minh hoạ chương Giới thiệu',
       'Adolescent girl — illustration for Introduction chapter')

H('Đối tượng chính (Key Audiences)', level=2)

P('Các nhà hoạch định chính sách giáo dục từ Bộ GD&ĐT Việt Nam là đối tượng chính của '
  'nghiên cứu và báo cáo này. Nghiên cứu này có thể cung cấp thông tin cho các nhà '
  'hoạch định chính sách giáo dục khi họ cân nhắc cách thức tốt nhất để hỗ trợ hạnh '
  'phúc học sinh vị thành niên. Tương tự, báo cáo cung cấp dữ liệu và khuyến nghị có '
  'khả năng có lợi cho các nhà hoạch định chính sách của MOH và MOLISA. Dữ liệu được '
  'thu thập bao gồm các lĩnh vực hỗ trợ của ngành y tế và xã hội cho sức khoẻ tâm thần '
  'vị thành niên, và các khuyến nghị bao gồm những đề xuất cụ thể cho từng ngành. Hơn '
  'nữa, các kết nối giữa các Bộ và liên ngành được khám phá trong nghiên cứu với các '
  'khuyến nghị cụ thể được đưa ra liên quan đến các lĩnh vực cần cải thiện hỗ trợ liên '
  'ngành. Nghiên cứu này cũng nhằm mục đích cung cấp thông tin cho UNICEF Việt Nam '
  'trong những nỗ lực vận động chính sách và tiến bộ chương trình trong lĩnh vực sức '
  'khoẻ tâm thần vị thành niên.')

P('Các nhà quản lý giáo dục, y tế và xã hội cấp tỉnh và huyện, hiệu trưởng, giáo '
  'viên, phụ huynh, học sinh cũng sẽ hưởng lợi từ các phát hiện và khuyến nghị chính '
  'của báo cáo này. Các nhà nghiên cứu, nhà hoạch định chính sách và nhà phát triển '
  'chương trình quốc tế làm việc trong lĩnh vực sức khoẻ tâm thần học sinh vị thành '
  'niên cũng có thể thấy nghiên cứu và báo cáo này hữu ích.')

H('Mục tiêu và Mục đích (Goals and Objectives)', level=2)

P('Nghiên cứu này nhằm xây dựng hiểu biết về việc các yếu tố liên quan đến trường học '
  'tác động như thế nào đến sức khoẻ tâm thần và hạnh phúc của thanh thiếu niên nam '
  'và nữ tại Việt Nam, cũng như vai trò của hệ thống giáo dục trong việc giải quyết cả '
  'các rủi ro MHPSS liên quan đến trường học và việc cung cấp dịch vụ. Nghiên cứu này '
  'bao gồm đánh giá mối quan hệ giữa các thách thức sức khoẻ tâm thần vị thành niên '
  '(lo âu, trầm cảm, stress) và các yếu tố nguy cơ liên quan đến trường học (khí hậu '
  'trường học, áp lực học tập và stress xã hội). Các yếu tố ngữ cảnh như tuổi, giới, '
  'tình trạng kinh tế xã hội (học sinh từ cả trường THPT công lập và tư thục được khảo '
  'sát để mở rộng phạm vi kinh tế) và vùng địa lý (bốn tỉnh Hà Nội, Điện Biên, Gia Lai '
  'và Đồng Tháp được bao gồm để đại diện cho các vùng khác nhau của Việt Nam và cả khu '
  'vực đô thị lẫn nông thôn). Các yếu tố hệ thống liên quan đến mối quan hệ giữa các '
  'yếu tố trường học và sức khoẻ tâm thần học sinh được nghiên cứu thông qua dữ liệu '
  'định tính thu thập từ mỗi cấp của hệ thống giáo dục (trường, huyện, tỉnh, bộ) và '
  'xuyên suốt các ngành chính phủ tác động đến sức khoẻ tâm thần học sinh (MOET, '
  'MOLISA, MOH).')

H('Câu hỏi nghiên cứu (Study Questions)', level=2)

P('Nghiên cứu giải quyết hai câu hỏi chính:')
P('1. Mối quan hệ giữa các yếu tố liên quan đến trường học và các vấn đề sức khoẻ tâm '
  'thần/hạnh phúc của thanh thiếu niên nam và nữ là gì; và')
P('2. Những can thiệp chương trình nào nhắm vào trường học có thể cải thiện sức khoẻ '
  'tâm thần và hạnh phúc vị thành niên cho nam và nữ tại Việt Nam.')

# --- Trang 9-10, UNICEF, 2022 — Study Limitations ---
PageMark('--- Trang 9–10, UNICEF Việt Nam, 2022 ---')

H('Hạn chế của nghiên cứu (Study Limitations)', level=2)

P('Các phát hiện của nghiên cứu bị hạn chế theo một số cách quan trọng. Nghiên cứu '
  'được hình thành và phát triển trong giai đoạn tháng 12/2020 – tháng 1/2021. Trong '
  'thời gian này, Việt Nam đã kiểm soát được COVID-19. Ngoài tác động đáng kể đến du '
  'lịch và các ngành khác phụ thuộc vào du lịch quốc tế, cuộc sống ở Việt Nam tương '
  'đối bình thường. Vì trọng tâm của nghiên cứu chủ yếu là sức khoẻ tâm thần vị thành '
  'niên và trường học, các công cụ và phương pháp nghiên cứu được phát triển phối hợp '
  'với Bộ GD&ĐT để giải quyết các câu hỏi then chốt và không bao gồm phân tích tác '
  'động của đại dịch lên sức khoẻ tâm thần học sinh. Sau đó, vào tháng 5/2021, một đợt '
  'bùng phát COVID-19 đáng kể đã xuất hiện ở miền Nam Việt Nam, dẫn đến phong toả và '
  'đóng cửa trường học toàn quốc. Trong khi dữ liệu nghiên cứu đã được thu thập ở ba '
  'tỉnh nông thôn hơn là Điện Biên, Gia Lai và Đồng Tháp, việc thu thập dữ liệu tại '
  'Thành phố Hồ Chí Minh (HCMC) và Hà Nội không thể tiến hành theo kế hoạch vào thời '
  'điểm đó. Việc thu thập dữ liệu được phép tiến hành tại Hà Nội vào tháng 9/2021, sử '
  'dụng khảo sát điện tử cho trẻ em và giáo viên, và phỏng vấn và thảo luận nhóm trực '
  'tuyến. Việc thu thập dữ liệu tại HCMC chưa bao giờ thực hiện được do quy mô của đợt '
  'bùng phát tại thành phố này. Các hạn chế tiềm năng của nghiên cứu liên quan đến '
  'những sự kiện này bao gồm:')

P('1. Tác động tiềm tàng của stress liên quan đến COVID-19 lên vị thành niên ở Hà Nội, '
  'có thể khác với tác động của COVID-19 lên người tham gia từ Điện Biên, Gia Lai và '
  'Đồng Tháp vì việc thu thập dữ liệu ở các tỉnh đó diễn ra trước đợt bùng phát. Nếu '
  'không có dữ liệu trước – sau bùng phát, chúng tôi không thể xác định liệu các báo '
  'cáo của người tham gia về triệu chứng hoặc nhu cầu sức khoẻ tâm thần có bị ảnh hưởng '
  'bởi đợt bùng phát ở Hà Nội hay không. Chúng tôi phải giả định rằng điều đó là đúng '
  'khi xét đến các nghiên cứu về tác động của stress liên quan đến COVID-19 lên vị '
  'thành niên trên toàn thế giới.')

P('2. Khả năng khái quát hoá các phát hiện nghiên cứu cho học sinh và trường học trên '
  'toàn quốc có thể bị giảm sút do mất dữ liệu HCMC. HCMC đại diện cho một khu vực địa '
  'lý duy nhất của đất nước về mặt vị trí và đô thị hoá. Khả năng khái quát hoá của '
  'nghiên cứu cho các trường học và học sinh đô thị được củng cố bởi sự tham gia của '
  'người tham gia từ Hà Nội, nhưng sẽ được tăng cường hơn nếu bao gồm HCMC trong '
  'nghiên cứu.')

P('3. Kế hoạch ban đầu là có hai cuộc phỏng vấn KII từ mỗi ba bộ, và hai cuộc phỏng '
  'vấn KII từ 3 sở xuyên suốt 5 tỉnh. Điều này đã đạt được ở ba tỉnh nông thôn nhưng '
  'đợt bùng phát tháng 5/2021 đã dẫn đến thách thức trong việc bảo đảm các cuộc phỏng '
  'vấn KII tại HCMC và Hà Nội. Ít nhất 1 KII đã được thực hiện từ mỗi Sở tại Hà Nội và '
  'từ mỗi Bộ. Việc có ít cuộc phỏng vấn hơn với các lãnh đạo ngành và chuyên gia có '
  'thể đã ảnh hưởng đến hiểu biết của chúng tôi về các thách thức chính sách và chương '
  'trình then chốt.')

P('Các phát hiện nghiên cứu cũng có thể bị hạn chế bởi các giả định nghiên cứu. Thiết '
  'kế nghiên cứu được xây dựng trên giả định rằng một số yếu tố liên quan đến trường '
  'học nhất định có khả năng cao nhất tác động đến sức khoẻ tâm thần và hạnh phúc học '
  'sinh. Các nghiên cứu trước đây chứng minh mối quan hệ giữa sức khoẻ tâm thần học '
  'sinh và các yếu tố trường học như khí hậu trường học, áp lực học tập và stress xã '
  'hội đã định hướng trọng tâm của chúng tôi vào những yếu tố này. Nghiên cứu nhằm '
  'tập trung rộng vào các vấn đề có thể liên quan đến sức khoẻ tâm thần vị thành niên, '
  'chẳng hạn như nghiên cứu các yếu tố gia đình như sự hiểu biết của phụ huynh về sức '
  'khoẻ tâm thần vị thành niên, áp lực học tập từ phụ huynh và sự kết nối của phụ '
  'huynh với nhà trường, và các yếu tố cộng đồng như tiếp cận dịch vụ y tế hoặc xã '
  'hội để hỗ trợ sức khoẻ tâm thần vị thành niên. Tuy nhiên, nghiên cứu không xem xét '
  'tác động của các yếu tố nguy cơ tiềm năng khác, bao gồm (1) các yếu tố cơ sở hạ '
  'tầng trường học như chương trình giảng dạy, sĩ số lớp và đào tạo giáo viên; (2) '
  'các yếu tố gia đình như sức khoẻ tâm thần và hạnh phúc của phụ huynh và chất lượng '
  'mối quan hệ cha mẹ – con; và (3) các yếu tố xã hội như hiểu biết sức khoẻ tâm thần '
  'cộng đồng kém, nhận thức cộng đồng về vai trò giới, và kỳ thị sức khoẻ tâm thần và '
  'phân biệt đối xử đối với bệnh tâm thần, hoặc dựa trên bản dạng giới hoặc xu hướng '
  'tính dục. Phải giả định rằng những yếu tố này và các yếu tố khác cũng quan trọng '
  'đối với sức khoẻ tâm thần học sinh và phải được xem xét trong cách tiếp cận toàn '
  'diện đối với sức khoẻ và hạnh phúc vị thành niên.')

P('Mặc dù không phải là hạn chế theo đúng nghĩa, những thay đổi cơ cấu gần đây trong '
  'Bộ GD&ĐT cần được theo dõi để đảm bảo tiến bộ liên tục hướng tới các mục tiêu MHPSS '
  'học sinh. Trước tháng 6/2021, Vụ Giáo dục Chính trị và Công tác Học sinh – Sinh viên '
  '(PESAD) trong Bộ GD&ĐT chịu trách nhiệm chính về sức khoẻ tâm thần học sinh, và '
  'PESAD là đơn vị hợp tác tích cực trong nghiên cứu hiện tại. Vào tháng 6/2021, trách '
  'nhiệm này đã được chuyển sang Vụ Giáo dục Thể chất (PED) trong Bộ GD&ĐT. Quá trình '
  'chuyển giao trách nhiệm sức khoẻ tâm thần học sinh sang PED đã diễn ra suôn sẻ, cho '
  'thấy PED sẽ cam kết như nhau đối với các vấn đề sức khoẻ tâm thần học sinh.')

doc.add_page_break()

# ============================================================
# CHUONG 2 — TONG QUAN TAI LIEU TOAN CAU
# ============================================================
H('CHƯƠNG 2: SỨC KHOẺ TÂM THẦN VỊ THÀNH NIÊN — TỔNG QUAN TÀI LIỆU TOÀN CẦU', level=1)
P('(Chapter 2 — Adolescent Mental Health: Global Literature Review)',
  italic=True, align='center', size=11)

# --- Trang 11 ---
PageMark('--- Trang 11, UNICEF Việt Nam, 2022 ---')
P('Vị thành niên là giai đoạn chuyển tiếp giữa tuổi thơ ấu và tuổi trưởng thành. '
  'Đây là khoảng thời gian có những thay đổi lớn về thể chất, trí tuệ, cảm xúc và xã '
  'hội trong cuộc đời của một con người. Quá trình chuyển tiếp này đem đến những cơ '
  'hội phát triển thú vị, cũng như những điểm dễ bị tổn thương cho người trẻ khi các '
  'em phấn đấu trở thành những người trưởng thành khoẻ mạnh, hạnh phúc và có năng '
  'suất. Việc đảm bảo vị thành niên có tất cả sự hỗ trợ các em cần trong giai đoạn '
  'then chốt này là thiết yếu.')

H('2.1. Sức khoẻ tâm thần vị thành niên (Adolescent Mental Health)', level=2)

P('Sức khoẻ tâm thần là nền tảng cho sự phát triển lành mạnh và thành công của vị '
  'thành niên. Sức khoẻ tâm thần bao gồm hạnh phúc cảm xúc, tâm lý và xã hội của chúng '
  'ta. Nó ảnh hưởng đến cách chúng ta suy nghĩ, cảm nhận và hành động. Nó cũng giúp '
  'xác định cách chúng ta xử lý stress, quan hệ với người khác và đưa ra lựa chọn. '
  'Sức khoẻ tâm thần tích cực cho phép vị thành niên nhận ra toàn bộ tiềm năng của '
  'mình, ứng phó với các stress cuộc sống, học tập và làm việc hiệu quả, hình thành '
  'và duy trì các mối quan hệ lành mạnh, đồng thời đóng góp có ý nghĩa cho gia đình '
  'và cộng đồng.')

P('Tuy nhiên, nhiều vị thành niên gặp các vấn đề sức khoẻ tâm thần, đe doạ nghiêm '
  'trọng đến sự phát triển lành mạnh của các em. Trên toàn cầu, các vấn đề sức khoẻ '
  'tâm thần đặt ra gánh nặng bệnh tật lớn đối với vị thành niên. Năm 2019, khoảng '
  '15 %, hay khoảng 1 trong 7 vị thành niên trên toàn thế giới gặp các vấn đề sức '
  'khoẻ tâm thần (Polanczyk, 2015). Điều này có nghĩa là ước tính 175 triệu thanh '
  'thiếu niên nam và nữ chịu đựng các vấn đề cảm xúc và hành vi suy yếu, khiến các em '
  'có nguy cơ có kết quả xã hội, học tập và sức khoẻ kém trong giai đoạn then chốt '
  'của tuổi trẻ. Nhận thức được tầm nghiêm trọng của vấn đề này đối với sức khoẻ và '
  'chức năng của cá nhân và xã hội, các Mục tiêu Phát triển Bền vững của Liên Hợp '
  'Quốc đã đưa sức khoẻ tâm thần thành mục tiêu chính, nhằm giảm 1/3 tử vong sớm do '
  'các bệnh không lây nhiễm đến năm 2030 thông qua phòng ngừa và điều trị, và thông '
  'qua thúc đẩy sức khoẻ tâm thần và hạnh phúc (Mục tiêu 3.4).')

H('2.2. Các vấn đề sức khoẻ tâm thần phổ biến ở tuổi vị thành niên '
  '(Common Mental Health Problems in Adolescence)', level=2)

P('Các vấn đề sức khoẻ tâm thần phổ biến ở tuổi vị thành niên bao gồm trầm cảm, '
  'lo âu và các rối loạn hành vi. Trầm cảm ở tuổi vị thành niên được đặc trưng bởi '
  'cảm giác buồn bã dai dẳng, mất hứng thú và năng lượng, dễ cáu kỉnh, cảm thấy tiêu '
  'cực và vô dụng. Trên toàn cầu, trầm cảm là nguyên nhân gây bệnh tật và tàn tật '
  'đứng thứ tư trong số vị thành niên tuổi 15–19 và đứng thứ mười lăm đối với nhóm '
  '10–14 tuổi. Các em gái và phụ nữ trẻ có khả năng mắc rối loạn trầm cảm và có nỗ '
  'lực tự hại cao gấp đến ba lần so với các em trai (WHO, 2014). Tất cả thanh thiếu '
  'niên có thể tâm trạng thay đổi hoặc cáu kỉnh vào những thời điểm nhất định, nhưng '
  'trầm cảm gây ra những thay đổi nghiêm trọng về cảm xúc, hành vi và nhận thức ở '
  'người trẻ, khiến các em rất khó tham gia vào trường học, duy trì mối quan hệ với '
  'bạn bè và gia đình, cũng như hình dung một tương lai tươi sáng. Nếu trầm cảm kéo '
  'dài không được điều trị, vị thành niên có nguy cơ bỏ học, thất nghiệp khi trưởng '
  'thành, lạm dụng chất, mang thai/làm cha mẹ sớm và trầm cảm ở tuổi trưởng thành '
  '(Clayborn, Varin, Colman, 2019). Các kết quả tâm lý – xã hội kém thường liên quan '
  'đến nhau và có thể dẫn đến sự lan truyền các khó khăn suốt đời.')

P('Trong những trường hợp nặng, trầm cảm có thể dẫn đến tự tử. Trên toàn cầu, tự tử '
  'là nguyên nhân tử vong đứng thứ tư ở nhóm 15–19 tuổi (WHO, tháng 6/2021). Ước tính '
  '62.000 vị thành niên đã tử vong vào năm 2016 do tự hại. 77 % các vụ tự tử toàn '
  'cầu xảy ra ở các quốc gia thu nhập thấp và trung bình. Mặc dù các em gái báo cáo '
  'tỷ lệ suy nghĩ tự tử cao hơn, nhiều nam giới tử vong do tự tử (12,6 trên 100.000 '
  'nam giới so với 5,4 trên 100.000 nữ giới).')

P('Các phương thức tự tử phổ biến bao gồm nuốt thuốc trừ sâu, treo cổ và súng. Các '
  'yếu tố nguy cơ của tự tử có nhiều mặt, bao gồm sử dụng rượu có hại, lạm dụng trong '
  'thời thơ ấu, kỳ thị đối với việc tìm kiếm trợ giúp, rào cản tiếp cận chăm sóc và '
  'tiếp cận phương tiện. Giao tiếp qua phương tiện truyền thông kỹ thuật số về hành '
  'vi tự tử là mối quan tâm mới xuất hiện đối với nhóm tuổi này. Tại châu Á, tự tử '
  'nằm trong số các nguyên nhân tử vong hàng đầu ở thanh thiếu niên, và thanh thiếu '
  'niên châu Á có tỷ lệ lưu hành cao về ý tưởng tự tử đã từng có (11,7 %) và nỗ lực '
  'tự tử (2,4 %) tại sáu quốc gia thành viên ASEAN (Blum, Sudhinaraset, & Emerson, '
  '2012; Peltzer, Yi, & Pengpid, 2017).')

# --- Trang 13 ---
PageMark('--- Trang 13, UNICEF Việt Nam, 2022 ---')
P('Lo âu là một vấn đề sức khoẻ tâm thần phổ biến khác ở tuổi vị thành niên. Lo âu '
  'được đặc trưng bởi những lo lắng và sợ hãi mãnh liệt, quá mức và dai dẳng. Tất cả '
  'vị thành niên có stress và lo lắng vào những thời điểm nhất định, nhưng người trẻ '
  'mắc rối loạn lo âu trải qua đau khổ nghiêm trọng về cảm xúc, thể chất và nhận thức, '
  'và gặp khó khăn trong việc tập trung vào những việc khác ngoài lo lắng hay sợ hãi '
  'của mình. Rối loạn lo âu là nguyên nhân gây bệnh tật và tàn tật đứng thứ chín đối '
  'với vị thành niên 15–19 tuổi và thứ sáu đối với nhóm 10–14 tuổi. Rối loạn lo âu '
  'cao hơn ở các em gái so với các em trai và phổ biến hơn ở vị thành niên lớn tuổi '
  'hơn. Vị thành niên mắc rối loạn lo âu có nguy cơ tăng cao đối với lo âu ở tuổi '
  'trưởng thành, trầm cảm, lạm dụng ma tuý và thành tích học tập kém khi trở thành '
  'người trưởng thành trẻ (Woodward & Fergusson, 2001).')

P('Nhiều vị thành niên cũng gặp các rối loạn hành vi. Các rối loạn hành vi bao gồm '
  'rối loạn tăng động giảm chú ý (ADHD), được đặc trưng bởi khó khăn trong việc chú ý, '
  'hoạt động quá mức và hành động không quan tâm đến hậu quả, và rối loạn ứng xử, '
  'được đặc trưng bởi các hành vi phá hoại hoặc thách thức. Mặc dù tất cả vị thành '
  'niên có thể bị phân tâm, bốc đồng và có hành vi liều lĩnh, những người mắc rối '
  'loạn hành vi thường gặp khó khăn đáng kể trong việc điều chỉnh sự chú ý, cảm xúc '
  'và xung động của mình. Hơn nữa, phản ứng tiêu cực từ bạn bè, giáo viên và cha mẹ '
  'đối với những hành vi này có thể dẫn đến lòng tự trọng thấp và gia tăng các hành '
  'vi có vấn đề, bẫy vị thành niên trong vòng luẩn quẩn của hành vi tiêu cực và kết '
  'quả xấu đi. Rối loạn hành vi là nguyên nhân gây tàn tật đứng thứ hai ở các em trai '
  'vị thành niên nhỏ tuổi 10–14 và nguyên nhân gây tàn tật đứng thứ ba ở các em gái '
  '10–14. Rối loạn hành vi cũng là nguyên nhân gây tàn tật đứng thứ năm ở các em '
  'trai 15–19 tuổi (trong khi không còn là nguyên nhân hàng đầu gây tàn tật ở các em '
  'gái trong nhóm tuổi này) (UNICEF Country Dashboard, 2019). Rối loạn hành vi đặt '
  'vị thành niên vào nguy cơ bỏ học, lạm dụng chất và hành vi tội phạm, và có thể dẫn '
  'đến nghịch cảnh về sức khoẻ tâm thần, gia đình, xã hội và kinh tế khi trưởng thành '
  '(Colman I, Murray J, Abbott RA, Maughan, 2009).')

H('2.3. Các yếu tố nguy cơ cho sức khoẻ tâm thần kém ở tuổi vị thành niên '
  '(Risk Factors for Poor Mental Health in Adolescence)', level=2)

P('Các yếu tố nguy cơ cho các vấn đề sức khoẻ tâm thần vị thành niên bao gồm di '
  'truyền, stress, sang chấn và bạo lực, các vấn đề về bản dạng, môi trường gia đình, '
  'đói nghèo và lạm dụng chất. Các yếu tố quyết định quan trọng khác bao gồm chất '
  'lượng cuộc sống gia đình và các mối quan hệ với bạn bè. Bằng chứng chỉ ra rằng kết '
  'quả sức khoẻ tâm thần được gây ra bởi sự kết hợp của các yếu tố này, và rõ ràng '
  'rằng vị thành niên càng tiếp xúc với nhiều yếu tố nguy cơ thì tác động tiềm năng '
  'lên sức khoẻ tâm thần của các em càng lớn.')

H('2.3.1. Yếu tố nguy cơ di truyền (Genetic risk factors)', level=3)

P('Các yếu tố di truyền, bao gồm điều hoà di truyền ngoại (epigenetic regulation — '
  'tương tác di truyền với các yếu tố môi trường) và đa hình di truyền (genetic '
  'polymorphisms — các thay đổi DNA), đặt con người vào nguy cơ mắc các vấn đề sức '
  'khoẻ tâm thần. Tuy nhiên, các dấu hiệu di truyền hay các thay đổi đơn lẻ sẽ không '
  'dẫn đến sự phát triển của một rối loạn tâm thần. Sự kết hợp của một hay nhiều đa '
  'hình cụ thể và các yếu tố môi trường nhất định có thể dẫn đến sự phát triển của '
  'một rối loạn tâm thần. Các rối loạn có cơ sở di truyền mạnh bao gồm rối loạn lưỡng '
  'cực, tâm thần phân liệt và rối loạn phổ tự kỷ. Không rối loạn tâm thần nào có '
  'nguy cơ di truyền 100 % và các yếu tố môi trường thường ảnh hưởng mạnh mẽ đến sự '
  'phát triển của một rối loạn cụ thể.')

H('2.3.2. Stress', level=3)

P('Stress là một phần bình thường của cuộc sống và thường vô hại (hoặc thậm chí có '
  'ích). Tuy nhiên, cường độ, tần suất và thời gian kéo dài của stress sẽ khác nhau '
  'đối với mỗi người trong mỗi tình huống. Khi con người bị stress nặng, những thay '
  'đổi hoá học thần kinh trong cơ thể có thể kích hoạt hệ thần kinh tự chủ, khiến cơ '
  'thể phản ứng như thể đang gặp nguy hiểm. Nhiều yếu tố có thể làm trải nghiệm '
  'stress tệ hơn, bao gồm hỗ trợ xã hội hạn chế, khó khăn trong điều tiết cảm xúc, '
  'khó khăn trong dung nạp sự không chắc chắn, thiếu tự tin, cảm thấy bất lực hay '
  'choáng ngợp, và trải nghiệm nhiều nguồn gây stress. Khi stress trở nên choáng '
  'ngợp và kéo dài, nguy cơ các vấn đề sức khoẻ tâm thần và các vấn đề y tế tăng lên. '
  'Stress dài hạn làm tăng nguy cơ các vấn đề sức khoẻ tâm thần như lo âu và trầm '
  'cảm, các vấn đề sử dụng chất, các vấn đề giấc ngủ, đau và các phàn nàn cơ thể như '
  'căng cơ.')

# --- Trang 14 ---
PageMark('--- Trang 14, UNICEF Việt Nam, 2022 ---')
P('Vị thành niên thường trải nghiệm stress liên quan đến áp lực học tập, xung đột '
  'gia đình, mong muốn có quyền tự chủ lớn hơn, áp lực phải phù hợp với bạn bè và '
  'phát triển tình dục. Ngày càng có nhiều bằng chứng cho thấy việc vị thành niên sử '
  'dụng công nghệ có thể đóng góp đáng kể vào stress vị thành niên qua bắt nạt mạng '
  '(cyberbullying) và cảm nhận khoảng cách giữa thực tế sống của vị thành niên và '
  'nhận thức của các em về trải nghiệm của người khác trên mạng. Có bằng chứng đáng '
  'kể rằng các sự kiện cuộc sống gây stress dự báo sự phát triển các vấn đề sức khoẻ '
  'tâm thần ở tuổi vị thành niên (Reiss và cộng sự, 2019). Mức độ stress cao ở tuổi '
  'vị thành niên đã được liên kết với nhiều hậu quả tiêu cực, bao gồm giảm thành tích '
  'học tập (Kaplan và cộng sự, 2005), bỏ học (Dupéré và cộng sự, 2015), gia tăng các '
  'vấn đề sức khoẻ tâm thần (Snyder và cộng sự, 2017), và giảm hạnh phúc (Chappel và '
  'cộng sự, 2014).')

H('2.3.3. Sang chấn (Trauma)', level=3)

P('Thuật ngữ "stress sang chấn" (traumatic stress) thường chỉ phản ứng thể chất và '
  'cảm xúc của một cá nhân đối với các sự kiện đe doạ tính mạng. Ở trẻ em và vị '
  'thành niên, các trải nghiệm có thể gây stress sang chấn bao gồm lạm dụng thể '
  'chất, lạm dụng tình dục, lạm dụng cảm xúc, bỏ bê, bạo lực giữa các cá nhân hoặc '
  'bị hại, bạo lực cộng đồng và thiên tai. Trẻ em và vị thành niên đặc biệt dễ bị '
  'tổn thương với bạo lực tình dục, có mối liên hệ rõ ràng với sức khoẻ tâm thần có '
  'hại. Tiếp xúc với các sự kiện sang chấn là phổ biến ở thanh thiếu niên, với '
  'khoảng hai phần ba trẻ em và vị thành niên trên toàn thế giới báo cáo tiếp xúc '
  'với ít nhất một sự kiện sang chấn (Saunders & Adams, 2014).')

P('Stress sang chấn có thể gây ra phản ứng thể chất và cảm xúc mãnh liệt, bao gồm '
  'cảm giác kinh hoàng, bất lực và kinh hoàng choáng ngợp, và một loạt cảm giác thể '
  'chất như tim đập mạnh, run rẩy, chóng mặt, buồn nôn, khô miệng và họng, và mất '
  'kiểm soát bàng quang hoặc ruột. Tiếp xúc với sang chấn có thể dẫn đến các vấn đề '
  'sức khoẻ tâm thần nghiêm trọng liên quan đến sang chấn, bao gồm rối loạn stress '
  'sau sang chấn (PTSD). Một phân tích tổng hợp các nghiên cứu công bố từ 1998 đến '
  '2011 cho thấy 15,9 % trẻ em và vị thành niên tiếp xúc với một sự kiện sang chấn '
  'đã phát triển PTSD [10]. Trẻ em từ các quốc gia thu nhập thấp và trung bình '
  '(LMICs) có thể có nguy cơ tiếp xúc với sang chấn và PTSD cao hơn; một nghiên cứu '
  'gần đây phát hiện khoảng một phần ba vị thành niên sống tại LMICs có một số triệu '
  'chứng PTSD sau khi trải qua một sự kiện sang chấn, trong khi gần một trong mười '
  'có đủ triệu chứng để chẩn đoán PTSD đầy đủ theo DSM-5. Vị thành niên có nguy cơ '
  'tiếp xúc với sang chấn cao hơn bao gồm các nhóm dân tộc thiểu số, thanh thiếu niên '
  'vô gia cư, trẻ em có cha mẹ có tiền sử hành vi tội phạm hoặc bệnh tâm thần, thanh '
  'thiếu niên LGBTQ và thanh thiếu niên tị nạn. Tiếp xúc với bạo lực (bao gồm nuôi '
  'dạy khắc nghiệt và bắt nạt) và các vấn đề kinh tế xã hội cũng được công nhận là '
  'rủi ro cho sức khoẻ tâm thần.')

P('Các trải nghiệm lạm dụng và bỏ bê bởi cha mẹ hoặc người chăm sóc trong thời kỳ '
  'ấu thơ đặt trẻ em vào nguy cơ cao mắc các vấn đề sức khoẻ tâm thần sau này trong '
  'đời. Sự gắn bó lành mạnh, đáp ứng với cha mẹ là cực kỳ quan trọng cho sự phát '
  'triển của trẻ. Khi cha mẹ liên tục đáp ứng các nhu cầu cơ bản của trẻ với sự nhạy '
  'cảm về cảm xúc, trẻ có thể lớn lên với cảm giác an toàn và có khả năng khám phá '
  'thế giới. Lạm dụng của cha mẹ, bao gồm lạm dụng cảm xúc, và bỏ bê gây stress "độc '
  'hại" đáng kể ở trẻ sơ sinh/trẻ đang phát triển, tác động nghiêm trọng đến sự phát '
  'triển nhận thức, xã hội và cảm xúc. Ngoài ra, hệ thống phản ứng stress sinh học '
  'có thể bị ảnh hưởng tiêu cực bởi lạm dụng và bỏ bê sớm, làm tăng nguy cơ lo âu, '
  'trầm cảm, lạm dụng chất và các suy giảm sức khoẻ mạn tính. Những trải nghiệm gây '
  'hại sớm này cũng liên quan đến kiểm soát xung động kém, tính tiêu cực cao, khó '
  'khăn học tập và thành tích học đường kém sau này (Lippard & Nemeroff, 2020).')

# --- Trang 15 ---
PageMark('--- Trang 15, UNICEF Việt Nam, 2022 ---')

H('2.3.4. Yếu tố nguy cơ và bảo vệ dựa trên gia đình '
  '(Family-Based Risk and Protective Factors)', level=3)

P('Các nghiên cứu từ khắp nơi trên thế giới đã phát hiện rằng chức năng gia đình '
  'tích cực và sự tham gia của cha mẹ vào cuộc sống của vị thành niên có liên quan '
  'đến giảm nguy cơ sức khoẻ tâm thần kém [19, 35, Chia và cộng sự, 2020], trong khi '
  'sự thiếu ấm áp của cha mẹ và kiểm soát quá mức của cha mẹ có liên quan đến một '
  'loạt các vấn đề tâm lý, bao gồm trầm cảm, hành vi tự tử và tự hại ở vị thành niên '
  '[36, 37, 38].')

P('Tại Việt Nam, sự tham gia, giám sát và quan tâm của cha mẹ đối với thời gian rảnh '
  'của vị thành niên có liên quan đáng kể đến khả năng giảm bắt nạt bị hại và các '
  'vấn đề sức khoẻ tâm thần ở vị thành niên. Tuy nhiên, mức độ kiểm soát cao của cha '
  'mẹ có liên quan đến khả năng vị thành niên bị tấn công thể chất, cô đơn và các vấn '
  'đề sức khoẻ tâm thần, bao gồm ý tưởng tự tử (Nguyen và cộng sự, 2019).')

H('2.3.5. Rủi ro môi trường (Environmental Risks)', level=3)

P('Một số vị thành niên có nguy cơ mắc các tình trạng sức khoẻ tâm thần cao hơn do '
  'hoàn cảnh sống, kỳ thị, phân biệt đối xử hoặc loại trừ, hoặc thiếu tiếp cận các '
  'hỗ trợ và dịch vụ chất lượng. Những nhóm này bao gồm vị thành niên sống trong các '
  'bối cảnh nhân đạo và mong manh; vị thành niên mắc bệnh mạn tính, vị thành niên '
  'mang thai, vị thành niên làm cha mẹ, hoặc những người trong hôn nhân sớm và/hoặc '
  'cưỡng ép; trẻ mồ côi; và vị thành niên từ các nhóm dân tộc thiểu số, bản dạng '
  'giới, xu hướng tính dục hoặc các nhóm bị phân biệt đối xử khác. Cũng có bằng '
  'chứng đang nổi lên rằng tiếp xúc sớm với ô nhiễm không khí đặt trẻ em vào nguy cơ '
  'cao mắc các vấn đề cảm xúc và hành vi (Salma, Ahmed, Gita và cộng sự, 2022).')

P('Đói nghèo, hay stress kinh tế gia đình, đã dự báo các vấn đề sức khoẻ tâm thần '
  'vị thành niên ở nhiều quốc gia, bao gồm Đức (Reiss, 2019), Trung Quốc (Xu và cộng '
  'sự, 2019) và Hoa Kỳ (Assing-Murray & Lebrun-Harris, 2020). Tình trạng dân tộc '
  'thiểu số có liên quan đến gia tăng nguy cơ các vấn đề sức khoẻ tâm thần (Bains, '
  'S. & Gutman, L., 2021) và các trải nghiệm phân biệt chủng tộc và phân biệt đối '
  'xử có tác động sâu sắc đến sức khoẻ tâm thần của trẻ em và vị thành niên dân tộc '
  'thiểu số (Trent, Dooley, & Dougé, 2019). Vị thành niên thiểu số tình dục có nguy '
  'cơ trầm cảm và các vấn đề sức khoẻ tâm thần khác cao hơn so với các bạn dị tính '
  '(Marshal MP, Dietz LJ, Friedman và cộng sự, 2011) và vị thành niên thiểu số tình '
  'dục có nguy cơ tự tử cao hơn (Luk và cộng sự, 2021).')

H('2.3.6. Yếu tố nguy cơ COVID-19 (COVID-19 Risk Factors)', level=3)

P('Đại dịch COVID-19 đã mang tính thảm hoạ, dẫn đến hơn 2 triệu ca tử vong trên '
  'toàn thế giới và khiến hàng tỷ người rơi vào cô lập xã hội. Các ước tính về vấn '
  'đề sức khoẻ tâm thần toàn cầu ở người trưởng thành chỉ ra rằng tỷ lệ rối loạn '
  'trầm cảm nặng đã tăng 27,6 % vào năm 2020, cho thấy thêm 53,2 triệu ca rối loạn '
  'trầm cảm nặng trên toàn cầu. Ước tính tỷ lệ rối loạn lo âu lan toả tăng 25,6 % '
  'trên toàn cầu trong năm 2020, phản ánh thêm 76,2 triệu ca trên toàn thế giới '
  '(Lancet COVID-19 Mental Disorders Collaborators, 2021).')

P('Vị thành niên có nguy cơ đặc biệt cao do đại dịch COVID-19. Stress liên tục về '
  'chức năng gia đình, bất an kinh tế, sợ lây nhiễm, cô lập xã hội và gián đoạn học '
  'tập đã có tác động tâm lý đáng kể lên sức khoẻ tâm thần vị thành niên. Bệnh tật '
  'và tử vong trong gia đình do COVID-19 đã gây stress sang chấn cho nhiều trẻ em '
  'và thanh thiếu niên trên toàn thế giới. Bạo lực trong gia đình, bao gồm bạo lực '
  'giữa vợ chồng và lạm dụng thể chất trẻ em, đã gia tăng trong đại dịch, đặt trẻ em '
  'và vị thành niên vào nguy cơ thể chất và tâm lý. Ngoài ra, vị thành niên đã dành '
  'nhiều thời gian hơn đáng kể trên mạng, nơi các em thường xuyên tiếp xúc với thông '
  'tin gây stress về đại dịch và với lạm dụng trực tuyến cùng bắt nạt mạng. Đây là '
  'tất cả những thách thức nghiêm trọng vì vị thành niên vẫn đang phát triển khả '
  'năng ứng phó và phục hồi, đặt người trẻ vào nguy cơ cao hơn mắc các vấn đề sức '
  'khoẻ tâm thần liên quan đến đại dịch. Và do các vấn đề sức khoẻ tâm thần ở tuổi '
  'vị thành niên đặt con người vào nguy cơ sức khoẻ tâm thần kém và kết quả sức '
  'khoẻ thể chất kém suốt đời, việc ưu tiên sức khoẻ tâm thần vị thành niên trong '
  'đại dịch là cực kỳ quan trọng.')

# --- Trang 16 ---
PageMark('--- Trang 16, UNICEF Việt Nam, 2022 ---')
P('Vị thành niên có vấn đề sức khoẻ tâm thần từ trước có nguy cơ cao hơn đối với '
  'các vấn đề sức khoẻ tâm thần liên quan đến đại dịch. Trải nghiệm cô lập, cảm giác '
  'không chắc chắn, thiếu thói quen hằng ngày, thiếu tiếp cận dịch vụ y tế cũng có '
  'thể đặt người trẻ vào nguy cơ khó khăn tâm lý – xã hội. Jones, Mitra & Bhuiyan '
  '(2021) đã thực hiện một tổng quan hệ thống 16 nghiên cứu định lượng được thực '
  'hiện trong 2019–2021 với 40.076 người tham gia. Trên toàn cầu, vị thành niên '
  'thuộc các nền tảng khác nhau đều trải qua tỷ lệ lo âu, trầm cảm và stress cao '
  'hơn do đại dịch. Vị thành niên cũng có tần suất sử dụng rượu và cần sa cao hơn '
  'trong đại dịch COVID-19. Tuy nhiên, hỗ trợ xã hội, kỹ năng ứng phó tích cực, '
  'cách ly tại nhà và thảo luận cha mẹ – con có vẻ như tác động tích cực đến sức '
  'khoẻ tâm thần vị thành niên trong giai đoạn khủng hoảng này.')

P('Một nghiên cứu định tính với vị thành niên tại Nha Trang và Vinh, Việt Nam đã '
  'tiết lộ các yếu tố ảnh hưởng đến sức khoẻ tâm thần và hạnh phúc trong đại dịch '
  'COVID-19. Các phát hiện cho thấy việc đóng cửa trường học và học trực tuyến có '
  'tác động lớn nhất đến hạnh phúc và thành tích học tập của học sinh. Các phát '
  'hiện tiếp tục chỉ ra rằng vị thành niên có nguy cơ tiếp xúc với nội dung internet '
  'có hại hoặc gây nghiện cao hơn do tăng thời gian dành trên mạng. Thanh thiếu '
  'niên cũng tiếp xúc với lượng lớn tin tức về COVID-19, làm tăng cảm giác sợ hãi '
  'và lo âu. Học sinh báo cáo cảm giác buồn bã và trầm cảm, cũng như lo âu xã hội '
  'sau thời gian xa bạn bè và giao lưu. Các em gái đặc biệt có nguy cơ về suy nghĩ '
  'liên quan đến hình ảnh cơ thể tiêu cực. Vị thành niên báo cáo rằng cha mẹ các em '
  'lo ngại về chất lượng học trực tuyến, và do đó đã thúc ép các em học tập mạnh '
  'hơn, thường dẫn đến căng thẳng gia đình (Samuels, Ho, Nguyen và cộng sự, 2021).')

H('2.4. Khoảng trống điều trị sức khoẻ tâm thần trẻ em và vị thành niên toàn cầu '
  '(Global Child and Adolescent Mental Health Treatment Gap)', level=2)

P('Tại các bối cảnh thu nhập thấp trên toàn cầu, chỉ một phần nhỏ những người chịu '
  'đựng sức khoẻ tâm thần kém sẽ có cơ hội tiếp cận hỗ trợ tâm thần hay tâm lý. Mặc '
  'dù các can thiệp hiệu quả để giảm triệu chứng sức khoẻ tâm thần đã tồn tại, nhiều '
  'người đau khổ không nhận được chăm sóc sức khoẻ tâm thần. Có đến 76–85 % người '
  'cần điều trị ở LMICs không nhận được bất kỳ điều trị nào (Patel và cộng sự, 2011; '
  'WHO, 2013). Bằng chứng chỉ ra rằng tình hình có thể còn tệ hơn; Patel phát hiện '
  'khoảng trống điều trị vượt quá 90 % đối với các rối loạn tâm thần phổ biến và rối '
  'loạn sử dụng rượu ở Ấn Độ và Trung Quốc, hai quốc gia thu nhập trung bình tương '
  'đối giàu nguồn lực (Fairburn & Patel, 2016). Việc thiếu chính sách sức khoẻ tâm '
  'thần, chương trình sức khoẻ tâm thần và luật pháp sức khoẻ tâm thần ở nhiều quốc '
  'gia, cũng như nguồn lực hạn chế (cả tài chính và nhân lực), cơ sở hạ tầng hạn chế, '
  'kỳ thị và xấu hổ là những lý do quan trọng cho việc tiếp nhận và phổ biến chăm '
  'sóc sức khoẻ tâm thần thấp (WHO, 2008; Munoz và cộng sự, 2016).')

P('Trẻ em và vị thành niên đối mặt với khoảng trống điều trị đặc biệt lớn (Morris '
  'và cộng sự, 2011). Sự thiếu hụt các dịch vụ sức khoẻ tâm thần trẻ em và vị thành '
  'niên (CAMH) đầy đủ trên toàn thế giới đã được báo cáo trong Bản đồ Sức khoẻ Tâm '
  'thần Trẻ em và Vị thành niên của WHO (WHO, 2005). Trên toàn thế giới, có rất ít '
  'bác sĩ tâm thần trẻ em và vị thành niên; tại các quốc gia thu nhập cao, số lượng '
  'bác sĩ tâm thần trẻ em là 1,19 trên 100.000 thanh thiếu niên, nhưng ở LMICs — nơi '
  'phần lớn trẻ em và vị thành niên thế giới sinh sống — con số là < 0,1 trên 100.000 '
  'dân (Skokauskas và cộng sự, 2019), tỷ lệ nhu cầu không được đáp ứng trong đào '
  'tạo CAMH cao hơn và khoảng cách lớn hơn giữa nhu cầu CAMH và dịch vụ CAMH có sẵn '
  'ở các quốc gia thu nhập thấp và trung bình so với các quốc gia thu nhập cao '
  '(Morris và cộng sự, 2011). Cũng đã được báo cáo rằng so với châu Âu và Bắc Mỹ, '
  'có những nhu cầu không được đáp ứng đáng kể về nguồn lực CAMH ở châu Á, trước số '
  'lượng trẻ em đang tăng nhanh cần đánh giá và chăm sóc sức khoẻ tâm thần liên tục.')

# --- Trang 17 ---
PageMark('--- Trang 17, UNICEF Việt Nam, 2022 ---')
P('Năm 2017, Việt Nam báo cáo chỉ có 900 bác sĩ tâm thần và chỉ 10 bác sĩ tâm thần '
  'chuyên về CAMH. Việt Nam thiếu Chính sách Quốc gia về Sức khoẻ Tâm thần Trẻ em '
  'và Vị thành niên và có nhu cầu tăng cường các bác sĩ tâm thần chuyên gia CAMH '
  'và các chuyên gia (Hirota, Guerrero, & Skokauskas, 2020). Cùng với ngành y tế, '
  'có một khoảng trống toàn cầu và cụ thể của Việt Nam trong hỗ trợ khu vực giáo '
  'dục và xã hội cho sức khoẻ tâm thần trẻ em và vị thành niên.')

H('2.5. Trường học và sức khoẻ tâm thần và hạnh phúc vị thành niên '
  '(Schools and Adolescent Mental Health and Well-being)', level=2)

P('Trường học đóng vai trò lớn trong cuộc sống của trẻ em và có tác động trực tiếp '
  'đến sức khoẻ tâm thần và hạnh phúc của các em. Môi trường trường học, trải nghiệm '
  'học tập và các mối quan hệ với giáo viên và bạn bè có thể vừa đặt trẻ em vào nguy '
  'cơ các vấn đề sức khoẻ tâm thần, vừa bảo vệ các em và thúc đẩy hạnh phúc tích '
  'cực. Với trường học là yếu tố quan trọng đối với sức khoẻ tâm thần và hạnh phúc '
  'của trẻ em, và với sức khoẻ tâm thần và hạnh phúc là yếu tố then chốt cho cả '
  'thành công học tập và phát triển, việc hiểu các mối liên kết giữa các yếu tố '
  'trường học, sức khoẻ tâm thần và hạnh phúc, và thành tích học tập là quan trọng.')

H('2.5.1. Các yếu tố nguy cơ sức khoẻ tâm thần dựa trên trường học '
  '(School-Based Mental Health Risk Factors)', level=3)

P('Một số yếu tố trường học có liên quan đến sức khoẻ tâm thần và hạnh phúc học '
  'sinh. Khí hậu trường học, áp lực học tập và các mối quan hệ bạn bè, bao gồm '
  'trải nghiệm bắt nạt, đều đã được phát hiện tác động đến sức khoẻ tâm thần học sinh.')

H('2.5.1.1. Khí hậu trường học (School Climate)', level=3)

P('Trong hai thập kỷ qua, khí hậu trường học đã nổi lên như một cấu trúc thống '
  'nhất tập trung vào nhiều trải nghiệm môi trường và trải nghiệm trường học, nhằm '
  'tìm hiểu rõ hơn cách những trải nghiệm này định hình kết quả học tập, hành vi và '
  'cảm xúc – xã hội của học sinh. Các nỗ lực để định nghĩa các khía cạnh then chốt '
  'của khí hậu trường học đã dẫn đến mô hình 3 yếu tố của khí hậu trường học, bao '
  'gồm 13 tiểu miền: an toàn (cảm nhận an toàn, bắt nạt và hung hăng, sử dụng ma '
  'tuý); gắn kết (kết nối với giáo viên, kết nối học sinh, gắn kết học tập, kết '
  'nối trường học, công bằng và tham gia của phụ huynh); và môi trường (quy tắc và '
  'hậu quả, tiện nghi thể chất và hỗ trợ, hỗn loạn) (Bradshaw, Waasdorp, Debnam, & '
  'Johnson, 2014). Một khía cạnh nghiêm trọng của khí hậu trường học là hình phạt '
  'thể xác, vẫn được sử dụng rộng rãi trong các trường học ở nhiều nơi trên thế '
  'giới. Trẻ em bị lạm dụng thể chất hoặc tình dục, quấy rối hoặc làm nhục bởi giáo '
  'viên thường phát triển các vấn đề cảm xúc và hành vi (Hinze, 2021). Stress tâm '
  'lý kết quả cũng tác động tiêu cực đến kỹ năng nhận thức và thành tích học tập '
  'của học sinh.')

P('Một phân tích tổng hợp gần đây của Wang và cộng sự (2020) đã khảo sát 61 '
  'nghiên cứu quốc tế được công bố từ 2000–2016 và phát hiện rằng khí hậu lớp học '
  'có liên quan tích cực với năng lực xã hội của học sinh, động lực và gắn kết, '
  'thành tích học tập, và liên quan tiêu cực với đau khổ cảm xúc – xã hội và các '
  'hành vi ngoại hoá (Wang, Degol, Amemiyaa, & Guoc, 2020). Một nghiên cứu được '
  'thực hiện tại Nam Úc phát hiện rằng nhận thức của học sinh về khí hậu trường '
  'học có liên quan tích cực với hạnh phúc của học sinh, khả năng phục hồi và bản '
  'dạng đạo đức (Riekie, Aldridge, & Afari, 2016). Xem xét những phát hiện này, '
  'ngày càng có nhiều quốc gia (ví dụ: Canada, Trung Quốc, Anh, Pháp, Đức, Israel, '
  'Singapore, Hoa Kỳ) tập trung vào cải thiện khí hậu lớp học như một mục tiêu '
  'trung tâm của các sáng kiến cải cách giáo dục nhằm thúc đẩy chất lượng trường '
  'học và hạnh phúc học tập và tâm lý của học sinh (Wang, Degol, Amemiyaa, & Guoc, '
  '2020).')

H('2.5.1.2. Áp lực học tập (Academic Pressure)', level=3)

P('Nghiên cứu chỉ ra rằng học sinh trên khắp thế giới thường báo cáo mức độ stress '
  'liên quan đến học tập cao (Pascoe, Hetrick & Parker, 2020). Tổ chức Hợp tác và '
  'Phát triển Kinh tế (OECD) đã thực hiện một khảo sát liên quan đến 72 quốc gia '
  'và bao gồm 540.000 học sinh tham gia tuổi 15–16 (Peña-López, 2016). Trung bình, '
  '66 % học sinh báo cáo cảm thấy stress về điểm kém và 59 % báo cáo rằng các em '
  'thường xuyên lo lắng rằng làm bài thi sẽ khó. OECD tiếp tục phát hiện rằng 55 % '
  'học sinh cảm thấy rất lo âu về thi cử trường học, ngay cả khi đã chuẩn bị tốt. '
  'Có tới 37 % học sinh báo cáo cảm thấy rất căng thẳng khi học, với các em gái '
  'liên tục báo cáo lo âu liên quan đến bài vở trường học lớn hơn so với các em '
  'trai. Dữ liệu này chứng minh rằng giáo dục và thành tích học tập là một nguồn '
  'stress đáng kể đối với học sinh.')

# --- Trang 18 ---
PageMark('--- Trang 18, UNICEF Việt Nam, 2022 ---')
P('Stress liên quan đến học tập được học sinh trung học trải qua tác động đến '
  'sức khoẻ tâm thần và thể chất của các em và dẫn đến một loạt các vấn đề học '
  'tập. Trong một nghiên cứu về học sinh vị thành niên tại Ấn Độ, vị thành niên có '
  'stress học tập có nguy cơ trầm cảm cao gấp 2,4 lần so với vị thành niên không có '
  'stress học tập. Tự kỳ vọng và áp lực cảm nhận từ cha mẹ và giáo viên là nguyên '
  'nhân chính của stress học tập. Các em gái vị thành niên có stress học tập cao '
  'hơn các em trai (Jayanthi, Thirunavukarasu, & Rajkumar, 2015). Nghiên cứu OECD '
  'được báo cáo ở trên phát hiện rằng lo âu về bài vở trường học, bài tập về nhà '
  'và thi cử có tác động tiêu cực đến thành tích học tập của học sinh trong khoa '
  'học, toán học và đọc hiểu (Peña-López, 2016).')

P('Stress liên quan đến học tập có liên quan mạnh mẽ đến giảm động lực học tập của '
  'học sinh và mất gắn kết học tập. Một nghiên cứu dọc trên 298 học sinh trung học '
  'Trung Quốc phát hiện rằng stress liên quan đến học tập ở Lớp 10 dự báo tiêu cực '
  'động lực học tập nội tại và dự báo tích cực việc thiếu động lực ở Lớp 12. Điều '
  'này chỉ ra rằng giảm stress liên quan đến học tập có thể bảo tồn động lực học '
  'tập nội tại liên tục của học sinh (Liu, 2015). Mối quan hệ giữa stress liên '
  'quan đến học tập, động lực và bỏ học dường như không phải là đặc thù văn hoá, '
  'với các phát hiện tương tự được chỉ ra từ nhiều nghiên cứu quốc tế (Pascoe, '
  'Hetrick & Parker, 2020; Walburg, 2014). Một nghiên cứu khác ở Trung Quốc đã '
  'khảo sát các mối quan hệ dọc giữa stress của vị thành niên ở trường và tỷ lệ '
  'thay đổi trong thành tích học tập của các em. Kết quả chỉ ra rằng đối với những '
  'học sinh có thành tích học tập giảm đáng kể theo thời gian, stress của học sinh '
  'từ tương tác giáo viên – học sinh đã dự báo đáng kể tỷ lệ thay đổi trong thành '
  'tích học tập của các em. Các phát hiện cung cấp hỗ trợ cho quan niệm rằng stress '
  'là yếu tố nguy cơ trong sự phát triển học tập của học sinh (Liu & Lu, 2011).')

H('2.5.1.3. Các mối quan hệ trong trường học (School Relationships)', level=3)

P('Các mối quan hệ lành mạnh, tích cực là then chốt cho sức khoẻ tâm thần và hạnh '
  'phúc học sinh. Thiếu tình bạn và bắt nạt là hai yếu tố đặt học sinh vị thành '
  'niên vào nguy cơ các vấn đề sức khoẻ tâm thần. Vị thành niên có khả năng có suy '
  'nghĩ tự tử cao hơn khi các em cảm thấy cô đơn. Các nghiên cứu trước đây được '
  'thực hiện ở Lebanon, Uganda, Tanzania, vùng cận Sahara châu Phi tiết lộ rằng trẻ '
  'em có nhiều khả năng có ý tưởng tự tử khi các em có cảm giác cô đơn (Mahfoud, '
  'Afifi, Haddad, & DeJong, 2011; Rudatsikira, Muula, Siziya, & Twa-Twa, 2007; '
  'Dunlavy, Aquah, & Wilson, 2015; Page & West, 2011). Học sinh vị thành niên ở '
  'Nepal trải qua sự cô đơn có nguy cơ có ý tưởng tự tử cao hơn, trong khi có 3 '
  'bạn thân trở lên được phát hiện là yếu tố bảo vệ (Pandey, Bista, Dhungana, '
  'Aryal, Chalise, & Dhimal, 2019).')

P('Bị bắt nạt và bắt nạt mạng (cyberbullying) là một yếu tố quan hệ khác đặt học '
  'sinh vào nguy cơ sức khoẻ tâm thần và hạnh phúc kém. Một nghiên cứu về học sinh '
  'tại Chile và Nam Phi phát hiện rằng bị bắt nạt là dự báo đáng kể của hạnh phúc '
  'trẻ em (Varela, Savahl, Adams & Reyes, 2019). Nghiên cứu chỉ ra rằng bắt nạt là '
  'yếu tố then chốt cho sức khoẻ tâm thần vị thành niên tại Việt Nam. Trong một '
  'nghiên cứu về bắt nạt và sức khoẻ tâm thần ở học sinh trung học và trung học '
  'phổ thông tại Việt Nam, bị bắt nạt đã dự báo độc lập các vấn đề sức khoẻ tâm '
  'thần của học sinh và các vấn đề sức khoẻ tâm thần dự báo trải nghiệm trở thành '
  'nạn nhân bắt nạt của học sinh. Các em gái có vấn đề sức khoẻ tâm thần có nhiều '
  'khả năng trở thành nạn nhân hơn; trong khi các em trai có vấn đề sức khoẻ tâm '
  'thần dễ bị cả bắt nạt lẫn trở thành thủ phạm (Le và cộng sự, 2019). Một nghiên '
  'cứu khác về học sinh vị thành niên tại Việt Nam phát hiện rằng bắt nạt phổ '
  'biến hơn ở trung học cơ sở so với trung học phổ thông (Nguyen, Nakamura, Seino '
  'và cộng sự, 2019).')

H('2.5.2. Sức khoẻ tâm thần và thành công học tập của học sinh '
  '(Student Mental Health and Academic Success)', level=3)

P('Ngoài việc làm suy yếu sức khoẻ và hạnh phúc tổng thể, các triệu chứng trầm '
  'cảm và lo âu còn có thể ảnh hưởng tiêu cực đến thành tích học tập (Bernal-'
  'Morales, Rodríguez-Landa, & Pulido-Criollo, 2015). Trầm cảm ở tuổi vị thành '
  'niên có tác động tiêu cực đáng kể lên thành tích học đường và đặt học sinh vào '
  'nguy cơ kết quả giáo dục và nghề nghiệp kém. Một nghiên cứu dọc 25 năm trên '
  'trẻ em New Zealand phát hiện những người bị trầm cảm ở tuổi 16–21 có tỷ lệ phụ '
  'thuộc phúc lợi và thất nghiệp cao hơn, chứng minh rằng tác động của sức khoẻ '
  'tâm thần kém ở tuổi vị thành niên có thể có các tác động kéo dài (Fergusson, '
  'Boden, & Horwood, 2007). Nghiên cứu tại Vương quốc Anh về học sinh lớp 7–9 phát '
  'hiện mối quan hệ mạnh giữa các triệu chứng trầm cảm của học sinh và điểm trung '
  'bình, với mối liên hệ mạnh hơn ở các em trai so với các em gái. Các triệu chứng '
  'trầm cảm cao hơn có liên quan đến thành tích học tập thấp hơn ở tuổi 16. Trầm '
  'cảm có liên quan đến khó khăn trong tập trung, quan hệ xã hội, thành tích học '
  'đường tự lực và đọc viết cũng như cảm nhận bài vở trường học là gánh nặng lớn. '
  'Nghiên cứu này chỉ ra rằng học sinh báo cáo khó khăn trong thành tích học tập '
  'nên được sàng lọc trầm cảm (Lopez-Lopez và cộng sự, 2021).')

P('Một số triệu chứng chính của trầm cảm, như khả năng tập trung suy giảm, mất '
  'hứng thú, thiếu chủ động, chậm chạp tâm vận động, tự trọng thấp, cảm giác vô '
  'dụng và rút lui xã hội, có thể làm rối loạn hoạt động nhận thức đáng kể và giảm '
  'sáng kiến trong học tập (Wagner, Müller, Helmreich và cộng sự, 2015). Trầm cảm '
  'có thể làm suy giảm chức năng nhận thức vì vị thành niên trầm cảm tập trung vào '
  'suy nghĩ và diễn giải trầm cảm thay vì nhiệm vụ thực tế, hoặc vì trầm cảm trực '
  'tiếp chặn các nguồn lực nhận thức, hoặc do cả hai lý do (Hartlage, Alloy, '
  'Vázquez, & Dykman, 1993).')

# --- Trang 19 ---
PageMark('--- Trang 19, UNICEF Việt Nam, 2022 ---')

H('2.5.3. Can thiệp dựa trên trường học (School-Based Interventions)', level=3)

P('Trường học cung cấp môi trường lý tưởng cho việc thúc đẩy sức khoẻ tâm thần và '
  'hạnh phúc vị thành niên. OECD nhấn mạnh rằng trường học là nơi người trẻ phát '
  'triển nhiều kỹ năng xã hội và cảm xúc cần thiết để trở nên kiên cường và thành '
  'công (OECD, 2015). Với hơn 66 % trẻ em ở độ tuổi trung học được ghi danh vào các '
  'trường trung học trên toàn thế giới năm 2018, các chương trình dựa trên trường '
  'học có thể tiếp cận nhiều nhóm vị thành niên và phụ huynh (UNESCO, tháng 2/'
  '2020). Hơn nữa, các dịch vụ sức khoẻ tâm thần dựa trên trường học đã được liên '
  'kết với kỳ thị thấp hơn và tỷ lệ sử dụng cao hơn, đặc biệt là trong số vị thành '
  'niên dân tộc thiểu số (Stephan và cộng sự, 2007). Do đó, các chương trình can '
  'thiệp dựa trên trường học cung cấp một môi trường đầy hứa hẹn cho chăm sóc '
  'phòng ngừa và tiếp cận thấp ngưỡng, với tiềm năng cũng tiếp cận các vị thành '
  'niên ngại tìm kiếm chăm sóc ngoài môi trường trường học hoặc sống trong cộng '
  'đồng không có lựa chọn chăm sóc khác. Vì sức khoẻ tâm thần và chức năng tâm lý '
  '– xã hội có lợi cho thành tích học tập và thành công trường học, các trường học '
  'có thể hưởng lợi từ việc triển khai các can thiệp nhằm cải thiện chức năng xã '
  'hội và cảm xúc.')

P('Một số chương trình sức khoẻ tâm thần và hạnh phúc vị thành niên dựa trên '
  'trường học đã được phát triển và đánh giá trên khắp thế giới. Các chương trình '
  'có xu hướng thuộc hai loại, bao gồm các chương trình phòng ngừa hoặc can thiệp '
  'phổ quát (tức là nhắm vào tất cả học sinh) và các can thiệp dựa trên trường học '
  'chọn lọc hoặc chỉ định (tức là nhắm vào học sinh có nguy cơ). Một số tổng quan '
  'quy mô lớn và phân tích tổng hợp đã khảo sát phạm vi và hiệu quả của các chương '
  'trình dựa trên trường học (Wells, Barlow, Stewart-Brown, 2003; van Loon và cộng '
  'sự, 2020; Greenberg, Domitrovich & Bumbarger, 2000; Hoagwood và cộng sự, 2007; '
  'Kraag, Zeegers, Kok, Hosman & Abu-Saad, 2006). Y văn nhấn mạnh các lợi ích học '
  'tập của việc thúc đẩy sức khoẻ tâm thần trong trường học và có bằng chứng mạnh '
  'mẽ cho hiệu quả của các can thiệp để hỗ trợ những cải thiện tích cực trong các '
  'kết quả xã hội – cảm xúc và học tập của học sinh (Fazel, Hoagwood, Stephan & '
  'Ford, 2014; Sancassiani và cộng sự, 2015).')

H('2.5.3.1. Các chương trình phổ quát dựa trên trường học '
  '(Universal School-Based Programmes)', level=3)

P('Các cách tiếp cận phổ quát đối với thúc đẩy sức khoẻ tâm thần dựa trên trường '
  'học bao gồm nhiều loại can thiệp. Một số nhằm thay đổi văn hoá của trường học, '
  'ví dụ cho phép giáo viên áp dụng cách tiếp cận khác đối với kỷ luật để đảm bảo '
  'các lớp học không bạo lực, hoặc để học sinh giảm hành vi bắt nạt. Những chương '
  'trình khác dựa trên lớp học và nhằm cung cấp một chương trình cụ thể cho tất cả '
  'trẻ em trong lớp. Một số chương trình đã kết hợp can thiệp toàn trường và dựa '
  'trên lớp học. Nội dung của các can thiệp thúc đẩy sức khoẻ tâm thần dựa trên '
  'trường học khác nhau rất nhiều, và các chương trình có thể nhằm phòng chống bạo '
  'lực, xây dựng kỹ năng nhận thức xã hội – cảm xúc, hoặc phòng chống tự tử. Các '
  'phương thức cung cấp bao gồm các kỹ thuật để thay đổi hành vi hoặc dạy kỹ năng, '
  'tham gia vào các hoạt động hợp tác hoặc giúp đỡ, đào tạo giáo viên hoặc phụ '
  'huynh, và thay đổi trong môi trường, hệ thống hoặc văn hoá trường học.')

# --- Trang 20 ---
PageMark('--- Trang 20, UNICEF Việt Nam, 2022 ---')
P('Các phương thức triển khai dao động từ chương trình giới hạn trong lớp học, '
  'đến những thay đổi liên tục trong trường học, đến sự tham gia của phụ huynh và '
  'cộng đồng (Wells, Barlow & Stewart-Brown, 2003). Chính sách và luật pháp nghiêm '
  'cấm việc sử dụng kỷ luật thể chất bởi giáo viên và nhân viên nhà trường là một '
  'can thiệp then chốt nhằm làm cho trường học và lớp học an toàn hơn cho học sinh.')

P('Các chương trình phổ quát dựa trên trường học đã chứng minh giảm stress học '
  'sinh hiệu quả với một số kết quả hỗn hợp. Kraag và cộng sự (2006) phát hiện rằng '
  'các can thiệp phổ quát cho dân số chung trong trường học đã giảm stress học '
  'sinh một cách hiệu quả. Tuy nhiên, Loon và cộng sự (2020) phát hiện rằng các '
  'chương trình can thiệp dựa trên trường học đã giảm stress học tập cho các tiểu '
  'nhóm học sinh. Sự khác biệt trong các phát hiện có thể được quy cho độ tuổi '
  'trẻ hơn của học sinh trong Kraag (9–14 tuổi) so với Loon (10–18 tuổi), với '
  'trẻ nhỏ tuổi hơn hưởng lợi nhiều hơn một cách phổ quát từ các can thiệp dựa '
  'trên trường học. Phân tích tổng hợp Kraag 19 thử nghiệm ngẫu nhiên có đối '
  'chứng hoặc nghiên cứu bán thực nghiệm phát hiện rằng các chương trình trường '
  'học nhắm vào quản lý stress hoặc kỹ năng ứng phó đã giảm triệu chứng stress '
  'và cải thiện kỹ năng ứng phó ở học sinh (Kraag, Zeegers, Kok, Hosman & Abu-'
  'Saad, 2006). Das, Salam, Lassi và cộng sự (2016) đã rà soát các phân tích '
  'tổng hợp các can thiệp sức khoẻ tâm thần vị thành niên dựa trên trường học. '
  'Các phát hiện chỉ ra rằng các yếu tố liên quan đến thành công chương trình '
  'tổng thể bao gồm trọng tâm vào thúc đẩy sức khoẻ tâm thần chứ không phải phòng '
  'ngừa bệnh tâm thần. Mặc dù có ít bằng chứng từ LMICs, nghiên cứu hiện có gợi ý '
  'rằng phần lớn các chương trình kỹ năng sống và khả năng phục hồi dựa trên '
  'trường học đã chỉ ra tác động tích cực lên lòng tự trọng, động lực và tự hiệu '
  'quả của học sinh (Das, Salam, Lassi và cộng sự, 2016).')

P('Dự án Tăng cường Cơ sở Bằng chứng về Can thiệp Dựa trên Trường học để Thúc đẩy '
  'Sức khoẻ Vị thành niên (SEHER) đã được nghiên cứu rộng rãi tại Bihar, Ấn Độ. '
  'Chương trình là một can thiệp thúc đẩy sức khoẻ đa thành phần bao gồm các thành '
  'phần tập trung vào toàn trường, lớp học và cá nhân để thúc đẩy khí hậu trường '
  'học và sức khoẻ cùng hạnh phúc của vị thành niên tại các trường trung học. '
  'Chương trình được đánh giá trong một thử nghiệm liên quan đến 74 trường học và '
  'được cung cấp bởi một tư vấn viên giáo dân hoặc bởi một giáo viên hiện có. Can '
  'thiệp bao gồm các thành phần tập trung vào toàn trường, lớp học và cá nhân. '
  'Học sinh ghi danh vào lớp 9 (13–15 tuổi) đã được tiếp xúc với can thiệp trong '
  'hai năm. Đánh giá 8 tháng sau chương trình cho thấy kết quả đáng kể khi chương '
  'trình được cung cấp bởi một tư vấn viên giáo dân về khí hậu trường học, trầm '
  'cảm, bắt nạt, thái độ đối với bình đẳng giới, bị bạo lực và thực hiện bạo lực. '
  'Cỡ hiệu ứng đối với các kết quả này vào cuối năm 2 lớn hơn so với cuối năm 1. '
  'Các phát hiện của nghiên cứu cung cấp bằng chứng mạnh mẽ cho các lợi ích của '
  'một can thiệp thúc đẩy sức khoẻ trường học đa thành phần, khi được cung cấp '
  'bởi tư vấn viên giáo dân, đối với khí hậu trường học và một loạt các kết quả '
  'sức khoẻ và hạnh phúc của vị thành niên. Không có hiệu quả nào được thấy khi '
  'cùng can thiệp được giáo viên cung cấp. Can thiệp SEHER đã được so sánh với '
  'Chương trình Giáo dục Vị thành niên (AEP) trong lớp học, dạy kỹ năng sống, '
  'nhưng không cho thấy bất kỳ hiệu quả nào (Shinde và cộng sự, 2020). Một phát '
  'hiện quan trọng là khí hậu trường học đã trung gian các hiệu quả của can thiệp '
  'trên tất cả ba kết quả quan tâm. Môi trường trường học nuôi dưỡng, được đặc '
  'trưng bởi các mối quan hệ hỗ trợ và gắn kết với giáo viên và bạn bè, cảm giác '
  'thuộc về và tham gia tích cực vào khí hậu trường học, dự báo tỷ lệ triệu chứng '
  'trầm cảm thấp hơn, trải nghiệm bắt nạt và thực hiện bạo lực thấp hơn.')

H('2.5.3.2. Các chương trình mẫu từ Việt Nam (Example Programmes from Viet Nam)',
  level=3)

P('Đặng và cộng sự (2016) đã thích ứng và đánh giá một chương trình sức khoẻ tâm '
  'thần và kỹ năng xã hội phổ quát dựa trên lớp học cho học sinh tiểu học. Chương '
  'trình, RECAP-VN, là một chương trình bán cấu trúc cung cấp cho học sinh đào tạo '
  'kỹ năng xã hội trong lớp học, và cho giáo viên tư vấn trong lớp học về triển '
  'khai chương trình và quản lý hành vi lớp học. Nghiên cứu phát hiện rằng chương '
  'trình có tác động tích cực lên cả kỹ năng xã hội và chức năng sức khoẻ tâm '
  'thần.')

# --- Trang 21 ---
PageMark('--- Trang 21, UNICEF Việt Nam, 2022 ---')
P('Các hiệu quả điều trị đáng kể được phát hiện trên cả kỹ năng xã hội (đối với '
  'học sinh nguy cơ thấp) và chức năng sức khoẻ tâm thần (đối với cả học sinh '
  'nguy cơ thấp và cao). Mặc dù chương trình này dành cho trẻ em nhỏ tuổi hơn, '
  'các kết quả cho thấy triển vọng cho việc triển khai hiệu quả các chương trình '
  'sức khoẻ tâm thần phổ quát dựa trên trường học tại Việt Nam.')

P('Đặng, Weiss, Lam, và Ho (2018) cũng đã thích ứng và đánh giá một chương trình '
  'giải quyết vấn đề cho trường trung học phổ thông được gọi là Viet Nam ACES '
  'ProS. Chương trình ACES ProS đã được thử nghiệm tại Đà Nẵng với 100 học sinh '
  'trung học Việt Nam. Các phát hiện chỉ ra rằng chương trình có tác động tích '
  'cực lên các vấn đề cảm xúc và hành vi được học sinh báo cáo. Các tác giả đã '
  'thảo luận các vấn đề về thích ứng văn hoá chương trình liên quan đến triển '
  'khai thành công.')

H('2.5.3.3. Các chương trình chọn lọc dựa trên trường học '
  '(Selective School-Based Programmes)', level=3)

P('Các chương trình can thiệp chọn lọc hoặc chỉ định dựa trên trường học được '
  'cung cấp cho học sinh được đánh giá có nguy cơ vì nhiều lý do khác nhau. Các '
  'chương trình này có thể bao gồm các lớp học toàn bộ học sinh đều được coi là '
  'có nguy cơ, hoặc có thể đưa các tiểu nhóm học sinh ra khỏi lớp học thường quy '
  'để tham gia chương trình. Để giảm stress và cải thiện hạnh phúc của vị thành '
  'niên, các chương trình này cung cấp các cách tiếp cận khác nhau và áp dụng các '
  'cơ chế thay đổi giả định khác nhau. Ví dụ, chánh niệm (mindfulness — tức là '
  'đưa sự chú ý không phán xét vào khoảnh khắc hiện tại thông qua các kỹ thuật '
  'thiền định và bài tập nhận thức), các bài tập thư giãn (ví dụ: thư giãn tiến '
  'triển, thư giãn cơ, thư giãn dựa trên hình dung) và đào tạo kỹ năng sống, bao '
  'gồm các kỹ thuật nhận thức – hành vi khác nhau (ví dụ: điều tiết cảm xúc, giải '
  'quyết vấn đề, giải quyết xung đột), thường được sử dụng.')

P('Bằng chứng từ các can thiệp dựa trên trường học gợi ý rằng các can thiệp nhóm '
  'nhắm đích và các chương trình trường học dựa trên liệu pháp nhận thức hành vi '
  '(CBT) được phát hiện là hiệu quả trong việc giảm triệu chứng trầm cảm và lo '
  'âu (Das, Salam, Lassi và cộng sự, 2016). Tuy nhiên, với phần lớn bằng chứng '
  'đến từ các quốc gia thu nhập cao (HICs), khả năng khái quát hoá của các phát '
  'hiện cho LMICs bị hạn chế. Michelson và cộng sự (2020) đã thử nghiệm một dịch '
  'vụ sức khoẻ tâm thần cho vị thành niên ở LMICs gọi là PRemIum for aDoLEscents '
  '(PRIDE) giữa học sinh 12–20 tuổi ở Ấn Độ. Các phát hiện hỗ trợ một cơ chế cung '
  'cấp chi phí thấp của một can thiệp giải quyết vấn đề cho các vấn đề sức khoẻ '
  'tâm thần phổ biến. Trong chương trình này, học sinh có các triệu chứng sức '
  'khoẻ tâm thần tăng cao đã được giới thiệu để hỗ trợ và nhận được các tập truyện '
  'tranh về giải quyết vấn đề và cách ứng phó với các khó khăn phổ biến. Chỉ riêng '
  'điều này đã dẫn đến giảm các triệu chứng sức khoẻ tâm thần tổng thể, bao gồm '
  'triệu chứng nội hoá và ngoại hoá, cải thiện chức năng và hạnh phúc. Học sinh '
  'sau đó nhận được hỗ trợ tư vấn giáo dân cho thấy giảm lớn hơn các vấn đề vị '
  'thành niên và stress cảm nhận. Đào tạo tư vấn viên bao gồm sổ tay bằng văn '
  'bản mô tả can thiệp giải quyết vấn đề, 5 ngày đào tạo tại văn phòng, 6 tuần '
  'thực hành có giám sát và 3 ngày đào tạo bổ sung theo khoảng thời gian hàng '
  'tháng trong khi thử nghiệm đang được tiến hành. Các tư vấn viên cũng tham gia '
  'các hoạt động nâng cao nhận thức nhằm thúc đẩy nhận thức và khả năng chấp nhận '
  'các hoạt động tư vấn giữa nhân viên nhà trường và học sinh.')

H('2.5.3.4. Các yếu tố thành công triển khai chương trình MHPSS dựa trên trường học '
  '(School-Based MHPSS Programme Implementation Success Factors)', level=3)

P('Bất kể một chương trình có thể thành công như thế nào khi được đánh giá, có '
  'nhiều lý do khiến các chương trình không thành công hoặc không còn được cung '
  'cấp theo thời gian. Những lo ngại về triển khai được minh hoạ bằng một nghiên '
  'cứu về chương trình phòng ngừa phổ quát dựa trên trường học FRIENDS tại Vương '
  'quốc Anh. Các kết quả nghiên cứu đã chứng minh hiệu quả của chương trình trong '
  'việc giảm lo âu và trầm cảm học sinh (Stallard, Simpson, Anderson, & Goddard, '
  '2008). Tuy nhiên, chương trình dựa vào sự đóng góp đáng kể từ các chuyên gia '
  'đào tạo, như nhà tâm lý học lâm sàng và tư vấn viên nhà trường, khiến chương '
  'trình khó triển khai và tốn kém như một phần thực hành thường quy trong các '
  'trường học. Theo dõi sau đó tiết lộ rằng phần lớn các trường tham gia đã không '
  'tiếp tục triển khai chương trình sau thử nghiệm bốn năm. Những phát hiện này '
  'khẳng định nhu cầu về các chương trình dựa trên trường học sáng tạo, ít tốn '
  'thời gian và nguồn lực, cũng như ít phụ thuộc vào giám sát và hỗ trợ bên '
  'ngoài để tuân thủ và triển khai.')

# --- Trang 22 ---
PageMark('--- Trang 22, UNICEF Việt Nam, 2022 ---')
P('Nghiên cứu chương trình sức khoẻ tâm thần dựa trên trường học chỉ ra các '
  'yếu tố triển khai chương trình làm tăng xác suất bền vững và chất lượng liên '
  'tục của chương trình. Những thành phần chương trình then chốt này bao gồm (a) '
  'triển khai chương trình nhất quán; (b) bao gồm phụ huynh, giáo viên hoặc bạn '
  'bè trong các chương trình; (c) sử dụng nhiều phương thức (ví dụ: kết hợp '
  'trình bày thông tin với đào tạo kỹ năng nhận thức và hành vi); (d) phản hồi, '
  'tư vấn và hỗ trợ giáo viên (ví dụ: các buổi đào tạo bồi dưỡng, quan sát lớp '
  'học, thảo luận nhóm nhỏ) và (e) tích hợp nội dung chương trình vào chương '
  'trình giảng dạy lớp học chung; và (f) các thành phần chương trình phù hợp '
  'với sự phát triển (Rones & Hoagwood, 2000). Các can thiệp thành công nhất có '
  'nhiều khả năng là các chương trình thúc đẩy sức khoẻ tâm thần được cung cấp '
  'liên tục trong thời gian dài, tức là một năm hoặc lâu hơn, và nhằm lôi cuốn '
  'mọi người trong trường bao gồm học sinh, nhân viên, gia đình và cộng đồng, '
  'và thay đổi môi trường và văn hoá của trường. Những cách tiếp cận này có '
  'thể đòi hỏi những thay đổi trong thái độ, niềm tin hoặc hành vi của giáo '
  'viên (Wells, Barlow, Stewart-Brown, 2003).')

H('2.5.3.5. Công nghệ và can thiệp sức khoẻ tâm thần vị thành niên '
  '(Technology and adolescent mental health interventions)', level=3)

P('Công nghệ đem đến giá trị tiềm năng cho việc làm cho các can thiệp sức khoẻ '
  'tâm thần vị thành niên chọn lọc sẵn sàng có cho học sinh ở LMICs. Các rào cản '
  'đối với việc tiếp nhận và phổ biến chăm sóc sức khoẻ tâm thần bao gồm thiếu '
  'chính sách và chương trình sức khoẻ tâm thần, nguồn lực hạn chế (cả tài chính '
  'và nhân lực), cơ sở hạ tầng hạn chế, kỳ thị và xấu hổ. Khi mức độ thâm nhập '
  'internet và quyền sở hữu điện thoại di động đặc biệt đang tăng trên toàn cầu, '
  'các công nghệ kỹ thuật số đem đến cơ hội vượt qua một số rào cản này. Sở hữu '
  'internet và điện thoại thông minh đang tăng trên toàn cầu với ước tính 81 % '
  'người dân (6,37 tỷ người dùng) sở hữu điện thoại thông minh vào năm 2021. Tại '
  'Việt Nam, ước tính có 63,1 % tỷ lệ thâm nhập điện thoại thông minh với khoảng '
  '61,37 triệu trong tổng số 97,34 triệu người dân sở hữu điện thoại thông minh '
  '(Pew Research Centre, 2019).')

P('Các ví dụ về công nghệ sức khoẻ tâm thần bao gồm MoodGYM, một chương trình CBT '
  'tương tác, trực tuyến, tự định hướng được thiết kế để phòng ngừa và/hoặc giảm '
  'triệu chứng lo âu và trầm cảm. Chương trình đã được hơn ba phần tư triệu người '
  'sử dụng kể từ năm 2001 và đã chứng minh giảm đáng kể lo âu và trầm cảm ở vị '
  'thành niên lên đến 17 tuổi (Calear, Christensen, Mackinnon, Griffiths, & '
  'O\'Kearney, 2009; Fairburn & Patel, 2016). Một chương trình CBT máy tính không '
  'hướng dẫn khác là TreadWill. Chương trình này được thiết kế để hoàn toàn tự '
  'động, hấp dẫn, dễ sử dụng và có thể tiếp cận ở LMICs. Đánh giá hiệu quả và mức '
  'độ gắn kết của TreadWill đã được thực hiện trong một thử nghiệm ngẫu nhiên có '
  'đối chứng mù đôi với 598 người trưởng thành tham gia ở Ấn Độ (Ghosh A, Cherian '
  'RJ, Wagle S và cộng sự, 2021). Việc sử dụng TreadWill đã giảm đáng kể các triệu '
  'chứng liên quan đến trầm cảm và lo âu. Nhìn chung, các chương trình chuyên '
  'biệt trực tuyến cho sức khoẻ tâm thần cho thấy triển vọng như là những can '
  'thiệp có thể mở rộng quy mô. Mối quan tâm chính là tỷ lệ tuân thủ thấp. Ví dụ, '
  'đánh giá TreadWill phát hiện 12,1 % người tham gia đã gắn kết với sử dụng '
  'trung bình và 5,5 % gắn kết với hoàn thành đầy đủ.')

P('Tại Việt Nam, Sobowale và cộng sự (2016) đã khám phá các nhận thức của thanh '
  'thiếu niên và phụ huynh Việt Nam đối với các can thiệp kỹ thuật số cho sức '
  'khoẻ tâm thần thanh thiếu niên như là bước đầu tiên để triển khai điều trị '
  'dựa trên internet tại Việt Nam. Các phát hiện gợi ý rằng các chương trình dựa '
  'trên internet cho sức khoẻ tâm thần thanh thiếu niên, đặc biệt là các can '
  'thiệp kết hợp các thành phần giáo dục tâm lý và kết nối mạng xã hội, sẽ được '
  'tiếp nhận tốt tại Việt Nam.')

doc.add_page_break()

# ============================================================
# CHUONG 3 — TONG QUAN TAI LIEU VE SKTT VTN TAI VN
# ============================================================
H('CHƯƠNG 3: TỔNG QUAN TÀI LIỆU VỀ SỨC KHOẺ TÂM THẦN VỊ THÀNH NIÊN TẠI VIỆT NAM',
  level=1)
P('(Chapter 3 — Adolescent Mental Health in Viet Nam Literature Review)',
  italic=True, align='center', size=11)

# --- Trang 24 ---
PageMark('--- Trang 24, UNICEF Việt Nam, 2022 ---')

H('3.1. Bối cảnh Việt Nam (Context of Viet Nam)', level=2)

P('Nước Cộng hoà Xã hội Chủ nghĩa Việt Nam có dân số trên 96 triệu người. Là một '
  'trong ba quốc gia Đông Nam Á hàng đầu có tăng trưởng GDP cao, Việt Nam đang trải '
  'qua những thay đổi kinh tế – xã hội nhanh chóng trong xã hội, như gia tăng di cư '
  'nội bộ từ nông thôn ra đô thị, thay đổi trong cơ cấu gia đình, và thay đổi trong '
  'vai trò của cha mẹ trong các gia đình hiện đại hoá. Mặc dù có những thay đổi '
  'này, phần lớn (65,6 %) 96 triệu người dân của đất nước vẫn tiếp tục sống ở khu '
  'vực nông thôn và phụ thuộc vào nông nghiệp tự cung tự cấp. Thu nhập bình quân '
  'hộ gia đình hàng năm trên đầu người là 2.236 USD vào năm 2019, với sự chênh lệch '
  'lớn giữa khu vực đô thị và nông thôn (2021 CEIC Data). Tỷ lệ ghi danh tiểu học '
  'thô cao ở mức 117 % tính đến năm 2020 (World Bank).')

P('Việt Nam có 54 dân tộc, với dân tộc Kinh chiếm phần lớn dân số (84 %). Một số '
  'dân tộc nhỏ, dưới 1.000 thành viên, và ngoại trừ nhóm Hoa (người Trung Quốc), '
  'các dân tộc thiểu số có tỷ lệ nghèo, ít học và thường sống ở những nơi hẻo lánh '
  'cao không cân xứng. Các thành viên của các nhóm dân tộc thiểu số chiếm 15 % dân '
  'số cả nước nhưng chiếm 70 % số người nghèo cực độ (đo bằng ranh giới nghèo cực '
  'độ quốc gia). Có bằng chứng đáng kể về bất bình đẳng sức khoẻ đối với các nhóm '
  'dân tộc thiểu số. Các chính sách và chương trình của chính phủ chưa được thích '
  'ứng và nhạy cảm với văn hoá, và có các báo cáo về thái độ xấu và phân biệt đối '
  'xử từ nhân viên y tế đối với người dân tộc thiểu số. Đồng thời, có những ví dụ '
  'về truyền thống và cấu trúc phụ quyền trong các nhóm dân tộc thiểu số cũng duy '
  'trì các hành vi sức khoẻ có hại (Badiani-Magnusson và cộng sự, 2012; Malqvist '
  'và cộng sự, 2013).')

H('3.2. Chăm sóc và điều trị sức khoẻ tâm thần tại Việt Nam '
  '(Mental Health Care and Treatment in Viet Nam)', level=2)

P('Ngành chăm sóc sức khoẻ tâm thần tại Việt Nam vẫn đang phát triển. Cho đến gần '
  'đây, chăm sóc và điều trị sức khoẻ tâm thần tại Việt Nam được cung cấp gần như '
  'hoàn toàn bởi ngành y tế (chủ yếu chỉ cho các bệnh tâm thần nặng nhất). Ngành '
  'giáo dục đã phát triển các chính sách tư vấn học đường bắt đầu từ năm 2005 và '
  'ngành xã hội đã áp dụng các chính sách sức khoẻ tâm thần cộng đồng bắt đầu từ '
  'năm 2011. Xem Chương 7 để biết thêm chi tiết về các chính sách và chương trình '
  'liên quan đến sức khoẻ tâm thần của từng Bộ. Chính phủ Việt Nam đã thành lập '
  'Chương trình Sức khoẻ Tâm thần Quốc gia (NMHP), tuy nhiên NMHP chỉ bao phủ '
  'khoảng 30 % cả nước và chỉ bao gồm một danh sách hạn chế các bệnh tâm thần. '
  'Trong khi chính phủ ước tính khoảng 15 % dân số cần các dịch vụ chăm sóc sức '
  'khoẻ tâm thần, nghiên cứu độc lập gợi ý rằng con số gần hơn với 20 đến 30 % '
  'dân số (Mah, 2018).')

P('Tổ chức Y tế Thế giới (WHO) đã khảo sát đất nước này vào năm 2014 và phát hiện '
  'chỉ có 0,91 bác sĩ tâm thần có sẵn trên 100.000 người. Số bác sĩ tâm thần có '
  'sẵn trên 100.000 dân tại Việt Nam tương đương với các nước láng giềng ASEAN: '
  'Malaysia có 0,76 bác sĩ tâm thần trên 100.000 người và Thái Lan có 0,87 bác sĩ '
  'tâm thần trên 100.000 người. Tuy nhiên, nó tụt xa sau các nền kinh tế phát triển '
  'như Singapore với 3,48 bác sĩ tâm thần trên 100.000 người, và Hoa Kỳ với 12,40 '
  'bác sĩ tâm thần có sẵn trên 100.000 người. Khoa Tâm thần của Đại học Y Hà Nội '
  'và Viện Sức khoẻ Tâm thần Quốc gia cung cấp các chương trình đào tạo bác sĩ tâm '
  'thần. Sinh viên y đa khoa có thể chọn chuyên khoa tâm thần một năm — mặc dù sự '
  'quan tâm vẫn thấp so với các ngành y khác. Có rất ít nhà cung cấp sức khoẻ tâm '
  'thần trẻ em và vị thành niên tại Việt Nam (Đặng & Weiss, 2012).')

P('Mạng lưới sức khoẻ tâm thần y khoa bao gồm các chương trình nội trú và điều trị '
  'ngoại trú cho chăm sóc tâm thần. Hệ thống bệnh viện tâm thần tại Việt Nam có 36 '
  'bệnh viện được thành lập trên khắp cả nước, với khoảng 6.000 giường. Hệ thống '
  'cung cấp dịch vụ thông qua mạng lưới bệnh viện nhà nước; có hai Bệnh viện Tâm '
  'thần Quốc gia, một ở miền Bắc tại Hà Nội và một ở thành phố Biên Hoà miền Nam. '
  '34 bệnh viện tâm thần cấp tỉnh còn lại được phân bố khắp cả nước.')

# --- Trang 25 ---
PageMark('--- Trang 25, UNICEF Việt Nam, 2022 ---')
P('Việc nhập viện đối với bệnh nhân trong tình trạng nặng, thường là tâm thần '
  'phân liệt, rối loạn lưỡng cực và động kinh, được chăm sóc tại các bệnh viện nội '
  'trú này. Khả năng tiếp cận cũng bị hạn chế bởi phân bố địa lý của các nhà cung '
  'cấp vì các bác sĩ tâm thần được bố trí tại các bệnh viện tâm thần, nơi chăm '
  'sóc sức khoẻ tâm thần tập trung. Việt Nam có khoảng 600 cơ sở chăm sóc sức '
  'khoẻ tâm thần ngoại trú có sẵn cho người dân địa phương đang tìm chăm sóc ngắn '
  'hạn. Các cơ sở ngoại trú này chủ yếu nằm ở các trung tâm đô thị lớn của đất nước.')

P('Bộ Y tế Việt Nam đã bắt đầu mở rộng tiếp cận cộng đồng thông qua các cơ sở chăm '
  'sóc ban đầu để tiếp cận những người mắc bệnh tâm thần. Tuy nhiên, hiệu quả của '
  'điều trị trong chăm sóc ban đầu bị hạn chế bởi sự thiếu đào tạo đầy đủ để sàng '
  'lọc các rối loạn tâm thần và sáng kiến mở rộng cộng đồng chủ yếu tập trung vào '
  'tâm thần phân liệt và động kinh, vốn không phải là các bệnh tâm thần phổ biến '
  'nhất ở thanh thiếu niên Việt Nam (Ng và cộng sự, 2011). Từ năm 2015, Bộ Y tế đã '
  'mở rộng chương trình mục tiêu để tiếp cận những người mắc lo âu và trầm cảm, và '
  'bệnh nhân nhi mắc tự kỷ và rối loạn tăng động giảm chú ý. Những nỗ lực này đã '
  'đối mặt với nhiều thách thức do sự thiếu đào tạo và kinh nghiệm của ngành y tế '
  'trong lĩnh vực sức khoẻ tâm thần trẻ em và vị thành niên (Cuong, 2017).')

P('Kỳ thị là một rào cản nghiêm trọng đối với chăm sóc và điều trị sức khoẻ tâm '
  'thần tại Việt Nam. Các rối loạn tâm thần đôi khi được gắn với sự xấu hổ và nhục '
  'nhã trong nước, điều này có thể ngăn cản các cá nhân nói cởi mở về nỗi đau của '
  'mình. Ví dụ, "bác sĩ tâm thần" dịch trực tiếp sang tiếng Anh là "doctors who '
  'treat madness" (bác sĩ chữa bệnh điên). Các rối loạn tâm thần thường được xem '
  'là dấu hiệu của yếu kém cá nhân hoặc gia đình. Phân biệt đối xử đối với người '
  'mắc bệnh tâm thần có thể dẫn đến khó khăn trong tìm việc làm, kết hôn hoặc thậm '
  'chí việc các thành viên gia đình kết hôn.')

P('Một nghiên cứu do dự án Young Lives có trụ sở tại Vương quốc Anh thực hiện đã '
  'phát hiện các yếu tố gây stress liên quan đến nghèo đói làm tổn hại đến sự phát '
  'triển lành mạnh của trẻ em (Thang & Hang, 2018). Xuyên suốt cả khu vực đô thị '
  'và nông thôn, thanh thiếu niên và phụ nữ dễ tổn thương thiếu nguồn lực đầy đủ '
  'cho nhu cầu của mình. Chăm sóc đặc biệt thiếu hụt đối với những người sống ở '
  'khu vực nông thôn không có phương tiện giao thông tiếp cận đến các thành phố '
  'lớn — như Hà Nội và Thành phố Hồ Chí Minh — nơi các dịch vụ có sẵn rộng rãi hơn. '
  'Tuy nhiên, số lượng chuyên gia chăm sóc sức khoẻ tâm thần hạn chế hiện đang làm '
  'việc trong lĩnh vực này không thể đáp ứng nhu cầu của các phân khúc dân số dễ '
  'tổn thương, đặc biệt khi nguồn lực có hạn.')

H('3.3. Các vấn đề sức khoẻ tâm thần vị thành niên tại Việt Nam '
  '(Adolescent Mental Health Problems in Viet Nam)', level=2)

H('3.3.1. Các nghiên cứu về trầm cảm, lo âu và các vấn đề hành vi '
  '(Studies of Depression, Anxiety and Behavioural Problems)', level=3)

P('Các nghiên cứu về sức khoẻ tâm thần vị thành niên tại Việt Nam chỉ ra rằng '
  'thanh thiếu niên Việt Nam trải qua gánh nặng đáng kể về bệnh tâm thần. Nghiên '
  'cứu về tỷ lệ lưu hành các vấn đề sức khoẻ tâm thần vị thành niên tại Việt Nam '
  'đã phát hiện tỷ lệ khác nhau. Weiss và cộng sự (2014) đã thực hiện một nghiên '
  'cứu dịch tễ học sức khoẻ tâm thần đại diện quốc gia trên 1.314 trẻ em từ 6 đến '
  '16 tuổi từ 60 địa điểm trên khắp Việt Nam và ước tính 12 % dân số trẻ em và vị '
  'thành niên (hơn 3 triệu người trẻ) có vấn đề sức khoẻ tâm thần cần dịch vụ. '
  'Nghiên cứu này phát hiện tỷ lệ các vấn đề cảm xúc ở các em gái cao hơn so với '
  'các em trai, và tỷ lệ các vấn đề hành vi ở các em trai cao hơn. Trong nghiên '
  'cứu này, tình trạng kinh tế xã hội cao hơn có liên quan đến tỷ lệ triệu chứng '
  'ADHD cao hơn.')

P('Trong một nghiên cứu về phong cách nuôi dạy và các vấn đề sức khoẻ tâm thần ở '
  'học sinh trung học phổ thông Việt Nam từ Hà Nội, Huế và TPHCM, 16,4 % trong '
  '757 người tham gia báo cáo các vấn đề sức khoẻ tâm thần. Các phát hiện cho '
  'thấy là nữ, đang học lớp 12, và có mẹ bảo vệ quá mức là các yếu tố nguy cơ đối '
  'với các vấn đề tâm thần, trong khi sự ấm áp của người cha làm giảm nguy cơ có '
  'vấn đề tâm thần ở vị thành niên (La và cộng sự, 2020).')

P('Một nghiên cứu trên 1.161 học sinh 15–19 tuổi đã khảo sát gánh nặng các vấn đề '
  'sức khoẻ tâm thần ở học sinh trung học tại thành phố Cần Thơ, Việt Nam. Tỷ lệ '
  'lưu hành các triệu chứng trầm cảm và lo âu có ý nghĩa lâm sàng lần lượt là '
  '41,1 % và 22,8 % (Nguyen và cộng sự, 2013).')

# --- Trang 26 ---
PageMark('--- Trang 26, UNICEF Việt Nam, 2022 ---')
AddImg('p026_img1_Im0.jpg',
       'Hình 3 (trang 26). Học sinh nữ dân tộc thiểu số — chương Việt Nam',
       'Ethnic minority schoolgirl — Viet Nam Literature chapter')
P('Học sinh nữ có xác suất có triệu chứng lo âu cao gấp ba lần so với học sinh '
  'nam. Trong một nghiên cứu khác về trầm cảm vị thành niên, Nguyễn và cộng sự '
  '(2013) phát hiện rằng 18,7 % trong 1.100 học sinh trung học có triệu chứng '
  'trầm cảm phù hợp với rối loạn trầm cảm nặng. Theo Khảo sát Đánh giá Thanh thiếu '
  'niên Việt Nam (SAVY I), 32 % thanh thiếu niên 14–25 tuổi báo cáo cảm thấy buồn '
  'về cuộc sống nói chung (Bộ Y tế, 2005).')

P('Một nghiên cứu trên hơn 4.500 thanh thiếu niên ở Hà Nội, bao gồm tỷ lệ cao '
  'người di cư, phát hiện tỷ lệ lưu hành 6 tháng của trầm cảm là 36 % (Nguyen và '
  'cộng sự, 2012). Trầm cảm có liên quan tích cực với uống rượu hiện tại: khả '
  'năng trầm cảm càng cao, người ta càng có khả năng uống rượu. Một yếu tố nguy '
  'cơ đối với trầm cảm trong nghiên cứu này là tình trạng di cư, với vị thành '
  'niên đã di cư đến Hà Nội cho thấy tỷ lệ trầm cảm cao hơn so với người không di '
  'cư. Tỷ lệ lưu hành khác nhau giữa các nghiên cứu có thể do các yếu tố phương '
  'pháp nghiên cứu, bao gồm cỡ mẫu, công cụ và phương pháp nghiên cứu khác nhau. '
  'Cũng có thể là các nghiên cứu chỉ xem xét vị thành niên ở các trung tâm đô thị '
  'lớn tìm thấy tỷ lệ các vấn đề sức khoẻ tâm thần cao hơn (ví dụ: Nguyen và cộng '
  'sự, 2012) so với các nghiên cứu bao gồm trẻ em sống ở khu vực nông thôn (ví dụ: '
  'Weiss và cộng sự, 2014). Dữ liệu lưu hành sức khoẻ tâm thần vị thành niên mạnh '
  'mẽ được kỳ vọng sẽ sớm có thông qua Khảo sát Sức khoẻ Tâm thần Vị thành niên '
  'Quốc gia (Erskine và cộng sự, 2021).')

H('3.3.2. Các nghiên cứu về nguy cơ tự tử (Studies of Suicide Risk)', level=3)

P('Điều rất đáng quan ngại là nguy cơ tự tử vị thành niên liên quan đến trầm cảm '
  'và các vấn đề sức khoẻ tâm thần khác. Các nghiên cứu về tự tử ở vị thành niên '
  'tại Việt Nam đã phát hiện tỷ lệ ý tưởng tự tử cao. Năm 2012–2013, Le và cộng '
  'sự (2016) đã khảo sát 1.745 học sinh 16–18 tuổi tại mười trường ở Hà Nội, hỏi '
  '"Trong 12 tháng qua, bạn có từng suy nghĩ nghiêm túc về việc cố tự tử không?" '
  'Kết quả cho thấy 21,4 % các em gái và 7,9 % các em trai báo cáo suy nghĩ tự '
  'tử trong 12 tháng trước. Kế hoạch tự tử trong năm trước được báo cáo bởi 7,8 % '
  'các em gái và 4,0 % các em trai. Học sinh trong nghiên cứu báo cáo tiếp xúc '
  'với nhiều hình thức bị hại, bao gồm lạm dụng và bỏ bê trẻ em, chứng kiến bạo '
  'lực gia đình hoặc khu phố, bị hại tài sản và bắt nạt mạng với 94,3 % đã trải '
  'qua ít nhất một hình thức bị hại trong đời và 31,1 % đã trải qua mười hình '
  'thức hoặc nhiều hơn.')

# --- Trang 27 ---
PageMark('--- Trang 27, UNICEF Việt Nam, 2022 ---')
P('Học sinh với mười trải nghiệm bị hại trở lên cho thấy tăng đáng kể các triệu '
  'chứng trầm cảm và lo âu, khả năng có hành vi nguy cơ sức khoẻ, hành vi tự tử '
  'và chất lượng cuộc sống liên quan đến sức khoẻ kém hơn ở cả các em gái và các '
  'em trai (Le và cộng sự, 2015; Le và cộng sự, 2016b).')

P('Nghiên cứu khác đã nghiên cứu tỷ lệ lưu hành nỗ lực tự tử ở vị thành niên. '
  'Thai (2010) đã nghiên cứu 1.226 vị thành niên và phát hiện rằng 5,8 % báo cáo '
  'cố tự tử trong 12 tháng qua. Một nghiên cứu chung (2015) giữa UNICEF, Viện '
  'Phát triển Hải ngoại (ODI), và Bộ Lao động, Thương binh và Xã hội (MOLISA) đã '
  'khảo sát khu vực nông thôn Điện Biên, miền Tây Bắc Việt Nam và ghi nhận 333 '
  'vụ cố tự tử, bao gồm 73 vụ tự tử hoàn tất. Trong 333 vụ cố tự tử, 140 vụ do '
  'trẻ em 19 tuổi hoặc nhỏ hơn thực hiện. Trong 73 vụ tự tử hoàn tất, 16 là trẻ '
  'em. Đáng báo động, cuộc khảo sát phát hiện rằng xu hướng tự tử là vấn đề cấp '
  'bách ở thanh thiếu niên Việt Nam trong tỉnh này. Các tác giả ghi nhận rằng sự '
  'sẵn có của lá cây độc ở Điện Biên dường như tạo điều kiện cho các nỗ lực tự '
  'tử, đặc biệt là ở các em gái người Hmong sống gần nơi lá cây độc mọc.')

H('3.3.3. Các nghiên cứu về yếu tố nguy cơ sức khoẻ tâm thần '
  '(Studies of Mental Health Risk Factors)', level=3)

P('Một yếu tố quyết định của các vấn đề sức khoẻ tâm thần vị thành niên là tiếp '
  'xúc với ngược đãi và lạm dụng. Nghiên cứu chỉ ra tỷ lệ cao ngược đãi và lạm '
  'dụng ở trẻ em Việt Nam. Khảo sát Chỉ số Mục tiêu Phát triển Bền vững về Trẻ em '
  'và Phụ nữ (SDGCW) gần đây của Việt Nam (2020–2021) về kỷ luật trẻ em phát hiện '
  '69 % trẻ em 10–14 tuổi đã trải qua kỷ luật bạo lực trong tháng trước, và các '
  'em trai có khả năng bị phạt bạo lực cao hơn các em gái. Một khảo sát trên '
  '2.591 vị thành niên Việt Nam 12–18 tuổi (Nguyen và cộng sự, 2010) phát hiện '
  'rằng 39 % báo cáo trải nghiệm lạm dụng cảm xúc, 47 % báo cáo trải nghiệm lạm '
  'dụng thể chất, gần 20 % báo cáo đã trải qua lạm dụng tình dục và 29 % báo cáo '
  'đã trải qua bỏ bê. Nghiên cứu này phát hiện tỷ lệ bỏ bê và lạm dụng cảm xúc '
  'cao hơn ở các em gái, và tỷ lệ lạm dụng thể chất cao hơn ở các em trai. Tương '
  'tự, Akmatov (2011) phát hiện rằng hơn 55 % vị thành niên tại Việt Nam đã trải '
  'qua lạm dụng thể chất vừa phải do cha mẹ gây ra, với 29 % vị thành niên báo '
  'cáo lạm dụng thể chất nặng. Nam giới trải qua tỷ lệ lạm dụng thể chất cao hơn '
  'trong nghiên cứu này. Các hộ gia đình lớn hơn và tình trạng kinh tế nghèo hơn '
  'có liên quan đến tỷ lệ lạm dụng cao hơn.')

P('Sự kết nối với cha mẹ và trường học cũng tác động đến sức khoẻ tâm thần vị '
  'thành niên. Một khảo sát gần đây về học sinh trung học phổ thông Việt Nam '
  '(Le và cộng sự, 2018) phát hiện giao tiếp kém giữa cha mẹ và vị thành niên có '
  'tác động tiêu cực lên sức khoẻ tâm thần của người trẻ, và góp phần vào lòng '
  'tự trọng thấp, cảm giác buồn và cô đơn và suy nghĩ tự tử. Cảm giác thuộc về '
  'và gắn kết với trường học của vị thành niên cũng đóng vai trò trong sức khoẻ '
  'tâm thần của các em. Dữ liệu từ Khảo sát Đánh giá Thanh thiếu niên Việt Nam '
  '(SAVY) I và SAVY II cho thấy vị thành niên cảm thấy kết nối với trường học '
  'của mình ít có khả năng báo cáo các triệu chứng tâm lý (Phuong và cộng sự, '
  '2013).')

P('Một yếu tố nguy cơ quan trọng khác đối với sức khoẻ tâm thần vị thành niên '
  'Việt Nam là sức khoẻ tâm thần của người chăm sóc. Stratton và cộng sự (2014) '
  'đã nghiên cứu tình trạng đau khổ tâm thần và sức khoẻ tổng quát của người '
  'chăm sóc (sử dụng Bảng câu hỏi Tự báo cáo-20) và các báo cáo về sức khoẻ tâm '
  'thần vị thành niên sử dụng phiên bản cha mẹ của Bảng câu hỏi Điểm mạnh và Khó '
  'khăn (SDQ). Sức khoẻ tâm thần người chăm sóc có liên quan tích cực với sức '
  'khoẻ tâm thần vị thành niên, với mối liên hệ có ý nghĩa sau khi kiểm soát các '
  'biến nhân khẩu học liên quan khác và tình trạng sức khoẻ tổng quát của người '
  'chăm sóc.')

P('Tổng thể, có vẻ như các yếu tố nguy cơ chính đối với sức khoẻ tâm thần kém của '
  'vị thành niên bao gồm giới tính nữ, tuổi vị thành niên lớn hơn, tình trạng di '
  'cư, sức khoẻ tâm thần người chăm sóc kém, giao tiếp cha mẹ – con kém, cảm giác '
  'không kết nối với trường học, và các trải nghiệm lạm dụng, sang chấn và bỏ bê.')

doc.add_page_break()

# ============================================================
# CHUONG 4 — PHUONG PHAP
# ============================================================
H('CHƯƠNG 4: PHƯƠNG PHÁP (Methodology)', level=1)

# --- Trang 29 ---
PageMark('--- Trang 29, UNICEF Việt Nam, 2022 ---')

P('Thiết kế phương pháp hỗn hợp (mixed-methods design) được sử dụng để điều tra '
  'các câu hỏi nghiên cứu. Nghiên cứu bao gồm tổng quan tài liệu và thu thập dữ liệu '
  'gốc định lượng và định tính. Thiết kế phương pháp hỗn hợp cho phép tích hợp dữ '
  'liệu một cách có hệ thống, xác nhận các phát hiện từ các nguồn thông tin khác '
  'nhau, và mở rộng kiến thức trong một số lĩnh vực then chốt. Các câu hỏi nghiên '
  'cứu then chốt được trả lời bằng dữ liệu tích hợp. Ví dụ, để hiểu rõ hơn về tỷ '
  'lệ lưu hành các vấn đề sức khoẻ tâm thần phổ biến ở học sinh vị thành niên tại '
  'Việt Nam, nghiên cứu của chúng tôi đã kết hợp (a) tổng quan tài liệu các phát '
  'hiện nghiên cứu trước đây, (b) dữ liệu định lượng gốc từ học sinh về các vấn '
  'đề sức khoẻ tâm thần tự báo cáo của các em, và (c) dữ liệu định tính về nhận '
  'thức của học sinh, giáo viên, phụ huynh và các nhà quản lý cấp trường, huyện '
  'và bộ về các vấn đề sức khoẻ tâm thần học sinh. Việc tích hợp các điểm dữ liệu '
  'này cung cấp hiểu biết vững chắc về sức khoẻ tâm thần học sinh vị thành niên.')

P('Một ví dụ về cách thiết kế phương pháp hỗn hợp cho phép xác nhận các phát hiện '
  'nghiên cứu có thể được tìm thấy trong các phân tích dữ liệu liên quan đến việc '
  'các nhà quản lý hệ thống chính phủ và trường học hiểu gì về các chính sách sức '
  'khoẻ tâm thần hiện có và các chương trình MHPSS. Bằng cách so sánh các chính '
  'sách và chương trình với hiểu biết và nhận thức của các bên liên quan khác '
  'nhau, chúng tôi có thể tìm hiểu về chính sách và chương trình nào là các phản '
  'ứng khả thi, hiệu quả đối với các vấn đề sức khoẻ tâm thần vị thành niên, và '
  'những chính sách và chương trình nào có thể cần được tăng cường hoặc sửa đổi '
  'để đáp ứng tốt hơn các nhu cầu của học sinh. Và cuối cùng, thiết kế phương pháp '
  'hỗn hợp cung cấp cơ hội để mở rộng kiến thức và hiểu biết của chúng tôi trong '
  'nhiều tình huống. Ví dụ, dữ liệu định lượng về mối quan hệ giữa các vấn đề sức '
  'khoẻ tâm thần học sinh và các chỉ số khí hậu trường học có thể gợi ý các lĩnh '
  'vực rủi ro sức khoẻ tâm thần và cung cấp thông tin cho các khuyến nghị. Tương '
  'tự, dữ liệu định tính về nhận thức của các bên liên quan về cách cải thiện sức '
  'khoẻ tâm thần học sinh đã mang lại một kho báu ý tưởng phong phú từ các bên '
  'liên quan có kinh nghiệm cao, những người hiểu các rào cản và cơ hội hệ thống, '
  'và do đó có thể đưa ra các ý tưởng khả thi với tiềm năng tác động tích cực cao.')

P('Dữ liệu định lượng được thu thập từ học sinh và giáo viên. Học sinh đã trả lời '
  'một khảo sát tự báo cáo bao gồm các thước đo về sức khoẻ tâm thần học sinh, '
  'nhận thức về khí hậu trường học, stress học tập, thời gian học tập, thành tích '
  'học tập và trải nghiệm bắt nạt mạng. Một khảo sát giáo viên cung cấp dữ liệu '
  'định lượng về nhận thức của giáo viên về sức khoẻ tâm thần và hạnh phúc học '
  'sinh, đào tạo giáo viên về sức khoẻ tâm thần học sinh, năng lực của giáo viên '
  'hỗ trợ sức khoẻ tâm thần và hạnh phúc học sinh, và nhận thức của giáo viên về '
  'các nguồn lực trường học cho học sinh cần hỗ trợ sức khoẻ tâm thần. Các phương '
  'pháp định tính bao gồm (1) các thảo luận nhóm tập trung (FGD) với học sinh, phụ '
  'huynh, giáo viên và (2) các phỏng vấn người cung cấp thông tin chủ chốt (KII) '
  'với hiệu trưởng, nhân viên cấp huyện từ Sở GD&ĐT, Sở Y tế và Sở Lao động, '
  'Thương binh và Xã hội cấp tỉnh, và các chuyên gia cấp Bộ từ Bộ GD&ĐT, Bộ Y tế '
  'và Bộ LĐ-TB&XH.')

H('4.1. Mục tiêu cụ thể, hoạt động và phương pháp nghiên cứu '
  '(Specific Aims, Activities and Research Methodology)', level=2)

P('Mục tiêu cụ thể 1: Cung cấp tổng quan các bằng chứng có sẵn và phân tích về '
  '(i) tình hình hiện tại của các vấn đề sức khoẻ tâm thần vị thành niên tại Việt '
  'Nam, (ii) các yếu tố liên quan đến trường học ảnh hưởng đến các vấn đề sức '
  'khoẻ tâm thần vị thành niên tại Việt Nam, và (iii) các bên liên quan chủ chốt, '
  'chính sách, luật pháp, tiêu chuẩn và chương trình liên quan đến sức khoẻ tâm '
  'thần và hạnh phúc vị thành niên tại Việt Nam.', bold=True)
P('Một cuộc tìm kiếm tài liệu được thực hiện bằng các cơ sở dữ liệu có liên quan '
  '(EBSCO, PubMed, psycINFO, các cơ sở dữ liệu VN). Tổng quan tài liệu bao gồm '
  'tổng quan về (a) dữ liệu dịch tễ học gần đây về các vấn đề sức khoẻ tâm thần '
  'phổ biến, (b) các yếu tố liên quan đến trường học tác động đến sức khoẻ tâm '
  'thần bao gồm khí hậu trường học (các yếu tố an toàn, gắn kết và môi trường), '
  'áp lực học tập và stress xã hội liên quan đến bạn bè, và (c) các dịch vụ và '
  'chương trình MHPSS dựa trên trường học hiện có tại Việt Nam. Rà soát các văn '
  'bản chính sách, văn bản pháp luật và các tiêu chuẩn, hướng dẫn liên quan từ '
  'các cơ quan chính phủ (Bộ GD&ĐT, Bộ LĐ-TB&XH và Bộ Y tế) đã được thực hiện '
  'để hiểu các khuôn khổ khái niệm độc lập và phụ thuộc lẫn nhau, vai trò, trách '
  'nhiệm và các dịch vụ hiện có được mỗi cơ quan cung cấp để hỗ trợ sức khoẻ tâm '
  'thần vị thành niên trong các trường học.')

# --- Trang 30 ---
PageMark('--- Trang 30, UNICEF Việt Nam, 2022 ---')

P('Mục tiêu cụ thể 2: Phân tích chuyên sâu các yếu tố nguy cơ then chốt liên quan '
  'đến trường học và tác động của chúng lên sức khoẻ tâm thần và hạnh phúc tâm lý '
  'của vị thành niên nam và nữ tại Việt Nam, bao gồm dữ liệu định tính và định '
  'lượng gốc.', bold=True)
P('Một cách tiếp cận nghiên cứu hệ thống (systems research) được sử dụng để giải '
  'quyết mục tiêu này. Nghiên cứu hệ thống dựa trên khái niệm rằng một hệ thống '
  'là một chức năng của các bộ phận, hoặc thành phần, của nó, và mỗi thành phần '
  'tương tác, kết nối, liên hệ và trong một số trường hợp, ảnh hưởng lẫn nhau. '
  'Vì vậy, những người từ mỗi cấp của hệ thống trường học tại Việt Nam đã được '
  'mời tham gia, bao gồm học sinh, phụ huynh, giáo viên, hiệu trưởng, cán bộ '
  'chính phủ cấp huyện và cán bộ chính phủ cấp bộ quốc gia. Xem phần Phương pháp '
  'Thu thập Dữ liệu Gốc bên dưới để biết thêm chi tiết.')

P('Mục tiêu cụ thể 3: Lập bản đồ các chính sách hiện có và các chiến lược, hành '
  'động của trường học trong việc nhận diện, giảm thiểu và giải quyết các vấn đề '
  'sức khoẻ tâm thần vị thành niên trong và ngoài trường học.', bold=True)
P('Bằng chứng từ tổng quan tài liệu, rà soát tài liệu chính sách, và các KII với '
  'cấp Bộ (Bộ Y tế, Bộ LĐ-TB&XH, Bộ GD&ĐT), cấp tỉnh và huyện (Sở Y tế, Sở LĐ-'
  'TB&XH và Sở GD&ĐT từ 5 tỉnh) và các hiệu trưởng của các trường trung học cơ '
  'sở và trung học phổ thông tham gia đã cung cấp thông tin cho việc lập bản đồ '
  'các chính sách, chiến lược và dịch vụ liên quan đến việc cung cấp hỗ trợ cấp '
  'một, cấp hai và cấp ba cho sức khoẻ tâm thần học sinh vị thành niên trong các '
  'trường học tại Việt Nam.')

P('Mục tiêu cụ thể 4: Khảo sát hiệu quả của các chính sách, hệ thống và nguồn '
  'nhân lực hiện có và các hạn chế của chúng.', bold=True)
P('Các chính sách liên quan đến hỗ trợ sức khoẻ tâm thần vị thành niên dựa trên '
  'trường học đã được so sánh với dữ liệu thu thập được về phát triển và triển '
  'khai các chương trình và dịch vụ. Các chính sách và hướng dẫn liên quan đến '
  'hợp tác liên cơ quan đã được so sánh với dữ liệu thu thập được về điều phối '
  'trường học địa phương, y tế và ngành xã hội trong chăm sóc và điều trị các vấn '
  'đề sức khoẻ tâm thần. Phân tích định tính dữ liệu KII đã xác định các chủ đề '
  'liên quan đến hiệu quả cảm nhận của các chính sách hiện có, các dịch vụ và '
  'điều phối hệ thống các dịch vụ.')

P('Mục tiêu cụ thể 5: Đề xuất các khuyến nghị và kế hoạch hành động về vai trò '
  'của hệ thống giáo dục trong việc cung cấp MHPSS và cách hệ thống giáo dục có '
  'thể giảm thiểu và phòng ngừa các rủi ro sức khoẻ tâm thần liên quan đến trường '
  'học để thúc đẩy sức khoẻ tâm thần tích cực cũng như phòng ngừa sức khoẻ tâm '
  'thần kém. Xem xét các phương thức khác nhau để tiếp cận vị thành niên nhỏ tuổi '
  'và lớn tuổi.', bold=True)
P('Các khuyến nghị được cung cấp dựa trên kết quả của các phân tích trên. Các '
  'yếu tố nguy cơ, các yếu tố bảo vệ, những khoảng trống trong chính sách và/hoặc '
  'triển khai chương trình MHPSS sẽ cung cấp thông tin cho các khuyến nghị với '
  'trọng tâm vào các chương trình MHPSS dựa trên trường học và các chiến lược '
  'toàn hệ thống có khả năng hiệu quả nhất đối với vị thành niên tại Việt Nam. '
  'Các khuyến nghị dựa trên bằng chứng được cung cấp để hướng dẫn mỗi cấp của hệ '
  'thống (Bộ, Tỉnh và Trường học) và mỗi ngành tác động đến sức khoẻ tâm thần vị '
  'thành niên, với trọng tâm chính vào ngành giáo dục và trọng tâm thứ cấp vào '
  'các ngành y tế và dịch vụ xã hội. Các khuyến nghị nhấn mạnh các cơ hội cho '
  'hợp tác cải thiện xuyên các cấp hệ thống và ngành.')

P('Mục tiêu cụ thể 6: Bao gồm dữ liệu và phân tích được phân tách từ các góc nhìn '
  'giới tính, và xem xét các chênh lệch về thu nhập, địa lý, dân tộc, xu hướng '
  'tính dục và bản dạng giới, tuổi và khuyết tật, bất cứ khi nào có thể áp dụng.',
  bold=True)
P('Các phân tích dữ liệu theo giới tính, tuổi, tình trạng kinh tế xã hội và khu '
  'vực địa lý được bao gồm. Học sinh dân tộc thiểu số được bao gồm (đặc biệt từ '
  'các trường ở các tỉnh Gia Lai và Điện Biên). Một nhóm nhỏ học sinh LGBTQ được '
  'bao gồm từ Hà Nội.')

H('4.2. Phương pháp thu thập dữ liệu gốc '
  '(Primary Data Collection Methodological Approach)', level=2)

H('4.2.1. Quy trình nghiên cứu hợp tác (Collaborative Research Process)', level=3)

P('Nhóm nghiên cứu đã làm việc chặt chẽ với UNICEF và Vụ Giáo dục Chính trị và '
  'Công tác Học sinh – Sinh viên của Bộ GD&ĐT để xác định các hoạt động và quy '
  'trình nghiên cứu cụ thể. Bộ GD&ĐT đóng vai trò then chốt trong việc lựa chọn '
  'các tỉnh phù hợp để đưa vào nghiên cứu (dựa trên các mục tiêu về đại diện địa '
  'lý, đại diện đô thị/nông thôn và sự tham gia của các cộng đồng dân tộc thiểu '
  'số).')

# --- Trang 31 ---
PageMark('--- Trang 31, UNICEF Việt Nam, 2022 ---')
P('Mặc dù đề xuất ban đầu không bao gồm Đồng Tháp, tỉnh này đã được đưa vào dựa '
  'trên đề xuất của Bộ GD&ĐT để bao gồm một tỉnh từ khu vực Đồng bằng sông Cửu '
  'Long. Bộ GD&ĐT cũng đã yêu cầu tăng cỡ mẫu nghiên cứu để tăng cường sức mạnh '
  'thống kê.')

P('Việc thu thập dữ liệu ở mỗi tỉnh được thực hiện bởi nhà nghiên cứu quốc gia '
  'cấp cao, TS. Đặng Hoàng Minh, và hai nhà tâm lý học cấp thạc sĩ được đào tạo '
  'về phương pháp nghiên cứu và thu thập dữ liệu. Ở mỗi tỉnh, nhóm nghiên cứu '
  'trước tiên gặp Sở GD&ĐT cấp tỉnh để rà soát các mục tiêu và mục đích nghiên '
  'cứu. Sở GD&ĐT đã làm việc với nhóm nghiên cứu để xác định các trường trung '
  'học cơ sở và trung học phổ thông tham gia, lên lịch thu thập dữ liệu dựa trên '
  'trường học (khảo sát, KII với Hiệu trưởng, và FGD), và lên lịch các KII với '
  'các sở giáo dục, xã hội và y tế cấp huyện. Việc thu thập dữ liệu ở mỗi tỉnh '
  'diễn ra trong khoảng thời gian 4 ngày (khoảng).')

H('4.2.2. Phương pháp thu thập dữ liệu (Data Collection Methods)', level=3)

P('Việc thu thập dữ liệu diễn ra tại các Bộ ở Hà Nội và tại bốn tỉnh đại diện '
  'cho các khu vực địa lý, văn hoá và đô thị/nông thôn khác nhau của đất nước, '
  'bao gồm Hà Nội, Đồng Tháp, Gia Lai và Điện Biên. Kế hoạch ban đầu bao gồm TPHCM '
  'phải được thay đổi do một đợt bùng phát COVID-19 nghiêm trọng đã ngăn cản các '
  'nỗ lực thu thập dữ liệu ở đó. Chúng tôi đã thu thập dữ liệu bằng các phương '
  'pháp sau: (a) phỏng vấn người cung cấp thông tin chủ chốt (KII) với các nhà '
  'quản lý cấp cao từ chính phủ và hệ thống trường học; (b) các thảo luận nhóm '
  'tập trung (FGD) với học sinh, phụ huynh và giáo viên; và (c) các khảo sát để '
  'thu thập dữ liệu định lượng từ học sinh và giáo viên.')

H('4.2.3. Người tham gia (Participants)', level=3)

P('1. Nhân sự then chốt cấp Bộ của chính phủ. KII với 1 nhân sự then chốt từ mỗi '
  'Bộ GD&ĐT, Bộ LĐ-TB&XH và Bộ Y tế, tổng cộng 3 KII với cán bộ chính phủ cấp Bộ.')
P('2. KII với Trưởng bộ phận Bảo vệ Trẻ em của UNICEF về các nghiên cứu trước đây '
  'liên quan và các tài liệu chính phủ then chốt.')
P('3. Nhân sự then chốt cấp tỉnh và cấp huyện của chính phủ. Ở các tỉnh Điện Biên, '
  'Đồng Tháp và Gia Lai, các nhân sự then chốt từ Sở GD&ĐT, Sở LĐ-TB&XH và Sở Y '
  'tế đã tham gia (2 người/cơ quan). Ở Hà Nội, do đợt bùng phát COVID-19, chỉ 1 '
  'nhân sự then chốt từ Sở GD&ĐT và Sở Y tế đã có thể tham gia (trong khi 2 từ '
  'Sở LĐ-TB&XH). Do đó, tổng cộng 22 KII với cán bộ chính phủ cấp tỉnh và huyện '
  'đã được bao gồm.')
P('4. Hiệu trưởng trường học. KII đã diễn ra với hiệu trưởng từ 2 trường ở mỗi '
  'tỉnh, bao gồm 1 trường trung học cơ sở và 1 trường trung học phổ thông. Ở Hà '
  'Nội, hai trường bao gồm một trường bán công. Điều này dẫn đến tổng cộng 7 KII '
  'với hiệu trưởng trường học.')

# Bảng người tham gia KII
P('Bảng 4.1. Số lượng người tham gia phỏng vấn KII theo địa bàn', bold=True, italic=True)
t = doc.add_table(rows=15, cols=3)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
set_grid(t, [3.0, 6.0, 3.0])
rows_data = [
    ('Địa bàn', 'Đơn vị', 'Số người'),
    ('Hà Nội', 'Hiệu trưởng / Sở Y tế / Sở LĐ-TB&XH / Sở GD&ĐT', '2 / 1 / 2 / 1'),
    ('Gia Lai', 'Hiệu trưởng (*) / Sở Y tế / Sở LĐ-TB&XH / Sở GD&ĐT', '2 / 2 / 2 / 2'),
    ('Điện Biên', 'Hiệu trưởng / Sở Y tế / Sở LĐ-TB&XH / Sở GD&ĐT', '2 / 2 / 2 / 2'),
    ('Đồng Tháp', 'Hiệu trưởng / Sở Y tế / Sở LĐ-TB&XH / Sở GD&ĐT', '2 / 2 / 2 / 2'),
    ('Cấp Bộ', 'Bộ GD&ĐT / Bộ Y tế / Bộ LĐ-TB&XH', '1 / 1 / 1'),
    ('Khác', 'UNICEF', '1'),
    ('', '', ''),
    ('', '', ''),
    ('', '', ''),
    ('', '', ''),
    ('', '', ''),
    ('', '', ''),
    ('', '', ''),
    ('TỔNG', '', '34'),
]
# Adjust rows
for ri, rd in enumerate(rows_data[:8] + rows_data[-1:]):
    if ri < len(t.rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri].cells[ci]; c.text = str(v)
            colw(c, [3.0, 6.0, 3.0][ci])
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                    if ri == 0 or (ri == 8 and rd[0] == 'TỔNG'):
                        r.bold = True
            if ri == 0: shade(c, 'D9E2F3')
# Remove extra rows
while len(t.rows) > 9:
    t._tbl.remove(t.rows[-1]._tr)

P('(*) Tại Gia Lai, một giáo viên chủ nhiệm đã tham gia phỏng vấn khi một hiệu '
  'trưởng không có mặt.', italic=True, size=10)

# --- Trang 32 ---
PageMark('--- Trang 32, UNICEF Việt Nam, 2022 ---')

P('5. Giáo viên. Dữ liệu khảo sát giáo viên được thu thập từ 6 giáo viên/trường, '
  'tổng cộng 66 khảo sát giáo viên (85 % nữ, tuổi trung bình 40, 61 % trường '
  'trung học cơ sở). Các FGD giáo viên ở ba tỉnh Điện Biên, Gia Lai và Đồng Tháp '
  'diễn ra từ tháng 4 đến tháng 5/2021. Hai FGD giáo viên diễn ra ở Hà Nội vào '
  'tháng 10/2021. Mỗi FGD bao gồm 7 giáo viên, tổng cộng 35 giáo viên tham gia '
  'trung học cơ sở và trung học phổ thông. Xem Phụ lục 3 để biết thông tin nhân '
  'khẩu học giáo viên chi tiết.')

P('6. Phụ huynh. Phụ huynh từ các tỉnh Điện Biên, Gia Lai và Đồng Tháp đã tham '
  'gia FGD vào tháng 4 và 5/2021. Bảy phụ huynh tham gia mỗi nhóm, tổng cộng 21 '
  'phụ huynh tham gia. FGD phụ huynh ở Hà Nội đã không diễn ra do các lo ngại '
  'của các trường về stress nuôi dạy con và các hạn chế COVID-19 trong đợt bùng '
  'phát COVID-19 vào thời điểm đó.')

P('7. Học sinh. Người tham gia học sinh bao gồm 668 học sinh (66 % nữ, tuổi '
  'trung bình 14,2, 54 % dân tộc thiểu số). Các FGD học sinh ở ba tỉnh Điện Biên, '
  'Gia Lai và Đồng Tháp diễn ra từ tháng 4 đến tháng 5/2021. Hai FGD học sinh '
  'diễn ra ở Hà Nội vào tháng 10/2021. FGD học sinh bao gồm 7 học sinh mỗi nhóm. '
  'Một FGD bổ sung được tổ chức cho một nhóm học sinh tự nhận là LGBTQ ở Hà Nội '
  'vào tháng 9/2021 (4 học sinh). Tổng cộng 39 học sinh đã tham gia các FGD cho '
  'dự án. Học sinh được bao gồm từ cả trường trung học cơ sở (học sinh 10–14 '
  'tuổi) và trường trung học phổ thông (học sinh 15–19 tuổi). Các FGD học sinh '
  'ở Gia Lai và Điện Biên bao gồm học sinh dân tộc thiểu số. Học sinh từ một '
  'trường bán công được bao gồm từ Hà Nội để mở rộng phạm vi kinh tế xã hội của '
  'người tham gia. Dữ liệu định lượng được thu thập qua các khảo sát từ khoảng '
  '60 học sinh/trường, tổng cộng 668 khảo sát học sinh. Xem Phụ lục 3 để biết '
  'thông tin nhân khẩu học học sinh chi tiết.')

H('4.3. Các công cụ đo khảo sát học sinh (Student Survey Measures)', level=2)

P('1. Sức khoẻ tâm thần học sinh được đánh giá bằng Bảng câu hỏi Điểm mạnh và '
  'Khó khăn – 25 (SDQ25), một bảng câu hỏi sàng lọc hành vi tự báo cáo ngắn gọn '
  'dành cho trẻ em 3–16 tuổi. Phiên bản tự báo cáo vị thành niên được sử dụng '
  'trong nghiên cứu phù hợp với người trẻ khoảng 11–16 tuổi (Goodman và cộng sự, '
  '1998). Các mục SDQ25 hỏi về đau khổ tâm lý không đặc hiệu bao gồm các triệu '
  'chứng trầm cảm, lo âu và các rối loạn dạng cơ thể. 25 mục này được chia giữa '
  '5 thang: các vấn đề cảm xúc (5 mục), các vấn đề hành vi (5 mục), tăng động/'
  'giảm chú ý (5 mục), các vấn đề quan hệ bạn bè (5 mục) và hành vi xã hội tích '
  'cực (5 mục). 20 mục từ 4 thang vấn đề tâm lý được cộng lại với nhau để tạo ra '
  'tổng điểm khó khăn. SDQ25 đã được sử dụng rộng rãi trên toàn thế giới và đã '
  'được thích ứng, xác thực và sử dụng rộng rãi tại Việt Nam (vd. Weiss và cộng '
  'sự, 2014).')

P('2. Khí hậu trường học được đánh giá bằng Khảo sát Khí hậu Trường học An toàn '
  'và Hỗ trợ Maryland được Thích ứng (Adapted Maryland Safe and Supportive '
  'Schools Climate Survey — MDS3) được phát triển bởi Trung tâm Phòng chống Bạo '
  'lực Thanh thiếu niên Johns Hopkins. MDS3 dành cho học sinh bao gồm 56 mục cốt '
  'lõi dựa trên các chỉ số an toàn, gắn kết và môi trường trường học đã được xác '
  'thực trước đây (Bradshaw, Waasdorp, Debnam & Johnson, 2014). Tất cả các lựa '
  'chọn trả lời được chấm trên thang Likert 4 điểm từ hoàn toàn đồng ý đến hoàn '
  'toàn không đồng ý, trong đó điểm cao hơn đại diện cho khí hậu trường học thuận '
  'lợi hơn. 56 mục đánh giá các lĩnh vực khí hậu trường học sau:')
P('a. An toàn (cảm nhận an toàn, bắt nạt và hung hăng, sử dụng ma tuý chung)')
P('b. Gắn kết (kết nối với giáo viên, kết nối học sinh, gắn kết học tập, kết nối '
  'toàn trường, văn hoá công bằng, tham gia của phụ huynh)')
P('c. Môi trường trường học (quy tắc và hậu quả, tiện nghi vật chất, hỗ trợ, hỗn '
  'loạn)')

# --- Trang 33 ---
PageMark('--- Trang 33, UNICEF Việt Nam, 2022 ---')
P('3. Stress học tập học sinh được đánh giá bằng Thang Stress Giáo dục cho Vị '
  'thành niên (Educational Stress Scale for Adolescents — ESSA) được phát triển '
  'và xác thực với vị thành niên Trung Quốc bởi Sun và cộng sự (2011). ESSA bao '
  'gồm 16 mục được học sinh đánh giá trên thang 5 điểm từ 1 (hoàn toàn không '
  'đồng ý) đến 5 (hoàn toàn đồng ý) với điểm cao hơn cho thấy stress lớn hơn. '
  'Thang đo bao gồm năm khía cạnh của stress giáo dục, gồm áp lực học tập (4 '
  'mục), lo lắng về điểm (3 mục), chán nản (3 mục), tự kỳ vọng (3 mục) và khối '
  'lượng công việc (3 mục). ESSA đã chứng minh độ nhất quán nội bộ tốt với '
  'Cronbach\'s alpha 0,81 cho thang chung và khoảng 0,66–0,75 cho mỗi khía cạnh '
  '(Sun và cộng sự, 2011). Trúc và cộng sự (2012) đã xác thực ESSA trong bối '
  'cảnh Việt Nam. Trong nghiên cứu Việt Nam, ESSA đã chứng minh độ nhất quán nội '
  'bộ thang chung tốt, độ hiệu lực yếu tố và độ hiệu lực đồng thời, tương tự như '
  'nghiên cứu gốc.')

P('4. Thời gian dành cho học sau giờ học. Học sinh được hỏi "Trung bình, bao '
  'nhiêu giờ thêm mỗi ngày bạn đã dành cho tự học sau giờ học trong học kỳ '
  'trước?" (Hầu như không có, dưới một giờ; 1–2 giờ; 2–3 giờ; và hơn 3 giờ). '
  'Thước đo này được sử dụng trong một nghiên cứu với mẫu Trung Quốc bởi Sun '
  '(2012) và mẫu Việt Nam bởi Thai (2010).')

P('5. Thời gian học thêm riêng. Người tham gia được yêu cầu nêu rõ thời gian '
  'thực tế (giờ) đã dành cho học thêm riêng một thầy một trò và học thêm riêng '
  'theo nhóm. Thời gian (giờ) dành cho học thêm riêng được đo bằng cách sử dụng '
  'câu hỏi: "Bạn đã dành bao nhiêu giờ mỗi tuần cho học thêm riêng một thầy một '
  'trò / học thêm riêng theo nhóm sau giờ học hoặc vào cuối tuần trong học kỳ '
  'trước?"')

P('6. Thành tích học tập. Người tham gia được yêu cầu nêu rõ điểm trung bình '
  'thực tế của 2 môn học chính và thành tích tổng thể (Toán, Ngữ Văn và điểm '
  'trung bình) trong học kỳ trước. Tại Việt Nam, thang điểm 10 đã được sử dụng '
  'xuyên suốt hệ thống trường học với 10 là cao nhất và 0 là thấp nhất (Bộ '
  'GD&ĐT Việt Nam, 2011). Học sinh được cung cấp các lựa chọn trả lời cho điểm '
  'dưới 5, 5–6,9, 7–7,9 và 8–10.')

P('7. Bắt nạt mạng (Cyberbullying) được đánh giá bằng một công cụ được thích ứng '
  'từ Bản kiểm kê Bắt nạt Mạng Sửa đổi (Revised Cyber Bullying Inventory — RCBI) '
  '(Topcu & Erdur-Baker, 2010) để sử dụng với vị thành niên Việt Nam (Trần, '
  'Nguyễn, Weiss, Nguyễn & Nguyễn, 2018). RCBI đã chứng minh độ nhất quán nội bộ '
  'đạt yêu cầu với Cronbach\'s alpha 0,82 và 0,75.')

P('Bản kiểm kê Bắt nạt Mạng Sửa đổi – Việt Nam bao gồm 18 mục về trải nghiệm '
  'của học sinh trên internet trong 6 tháng trước, bao gồm tần suất trải nghiệm '
  'và đau khổ về các trải nghiệm. Điểm cao hơn trên thang đo phản ánh tần suất '
  'bị bắt nạt mạng lớn hơn và stress liên quan lớn hơn.')

H('4.4. Khảo sát giáo viên (Teacher Survey)', level=2)

P('Các giáo viên trung học cơ sở và trung học phổ thông (N = 66) từ các tỉnh '
  'tham gia đã hoàn thành một khảo sát được phát triển cho dự án này, đánh giá '
  'nhận thức của giáo viên về sức khoẻ tâm thần và hạnh phúc học sinh, kinh '
  'nghiệm đào tạo giáo viên về sức khoẻ tâm thần học sinh, năng lực của giáo '
  'viên trong hỗ trợ sức khoẻ tâm thần và hạnh phúc học sinh, và nhận thức của '
  'giáo viên về các nguồn lực trường học cho học sinh cần hỗ trợ sức khoẻ tâm '
  'thần. Khảo sát mất 15–30 phút để hoàn thành.')

H('4.5. Thảo luận nhóm tập trung học sinh, phụ huynh, giáo viên '
  '(Student, Parent, Teacher Focus Group Discussions)', level=2)

P('Các FGD đã khám phá nhận thức về nhu cầu sức khoẻ tâm thần học sinh, trải '
  'nghiệm với các yếu tố nguy cơ liên quan đến trường học và nhận thức về các '
  'dịch vụ và can thiệp hiện có. Xem Phụ lục 1 để biết Hướng dẫn và Câu hỏi '
  'FGD. Các FGD kéo dài khoảng 60 phút mỗi cuộc và bao gồm các người tham gia sau.')

# --- Trang 34 ---
PageMark('--- Trang 34, UNICEF Việt Nam, 2022 ---')

H('4.6. Phỏng vấn người cung cấp thông tin chủ chốt — Hiệu trưởng, cấp Huyện và '
  'cấp Bộ Quốc gia (Principals, District and National Government KII)', level=2)

P('Các cuộc phỏng vấn KII kéo dài 45–60 phút mỗi cuộc, đã khám phá nhận thức về '
  'nhu cầu sức khoẻ tâm thần học sinh và các yếu tố nguy cơ liên quan đến trường '
  'học, hiểu biết về các chính sách liên quan và kiến thức, hiểu biết về các '
  'chương trình MHPSS hiện có. Xem Phụ lục 2 để biết Hướng dẫn và Câu hỏi KII.')

H('4.7. Phân tích dữ liệu (Data analysis)', level=2)

P('Dữ liệu định tính, bao gồm tất cả các FGD và KII, trước tiên đã được gỡ băng '
  'và sau đó được dịch sang tiếng Anh để phân tích bởi tác giả chính. Rà soát sơ '
  'bộ các bản gỡ băng đã được thực hiện để tìm kiếm các chủ đề và danh mục rộng. '
  'Dữ liệu đã được nhập vào Nvivo 12 để phân tích và được mã hoá theo 24 mã (bao '
  'gồm 32 mã phụ). Các mã này đã được tổ chức thành các chủ đề và chủ đề phụ. '
  'Dữ liệu định lượng, bao gồm dữ liệu khảo sát học sinh và giáo viên, đã được '
  'phân tích bằng SPSS 28. Dữ liệu mô tả đã được phân tích cho dữ liệu nhân khẩu '
  'học và dữ liệu biến then chốt (ví dụ: sức khoẻ tâm thần học sinh). Tương quan '
  'Pearson đã khảo sát các mối liên hệ giữa các biến then chốt (ví dụ: sức khoẻ '
  'tâm thần học sinh và các biến khí hậu trường học; sức khoẻ tâm thần học sinh '
  'và các biến áp lực học tập). Các kiểm định t đã khảo sát sự khác biệt trung '
  'bình trong các biến then chốt theo giới và cấp học (ví dụ: áp lực học tập '
  'theo giới). Phân tích phương sai (ANOVA) đã khảo sát sự khác biệt trong các '
  'biến then chốt theo tỉnh (ví dụ: 3 thang phụ khí hậu trường học theo 4 tỉnh).')

H('4.8. Các vấn đề đạo đức (Ethical Concerns)', level=2)

P('Các vấn đề đạo đức bao gồm bản chất nhạy cảm của nội dung dữ liệu định tính '
  'và định lượng. Các khảo sát học sinh bao gồm thông tin về sức khoẻ tâm thần '
  'học sinh, có thể có hậu quả tiêu cực cho trẻ em nếu được chia sẻ. Các khảo '
  'sát giáo viên bao gồm nhận thức về hỗ trợ trường học đối với học sinh có các '
  'vấn đề xã hội – cảm xúc, có thể làm không hài lòng các cấp trên ở trường hoặc '
  'ở các cấp cao hơn. Dữ liệu định tính bao gồm thông tin được chia sẻ về các '
  'trải nghiệm cá nhân của người tham gia với các vấn đề sức khoẻ tâm thần, và '
  'các lo ngại về phản ứng của trường học hoặc hệ thống đối với sức khoẻ tâm '
  'thần và hạnh phúc học sinh. Vì những lý do này, các khảo sát đã được phi định '
  'danh và được giữ tại một địa điểm an toàn. Các bản ghi âm đã được gỡ băng và '
  'phi định danh. Chỉ các thành viên nhóm nghiên cứu mới có quyền truy cập vào '
  'dữ liệu.')

P('Một lo ngại thứ hai là học sinh có thể cảm thấy bị áp lực phải tham gia vì '
  'việc thu thập dữ liệu diễn ra tại trường và có sự hỗ trợ của nhà quản lý '
  'trường. Để giảm thiểu rủi ro này, học sinh đã được thông báo rằng sự tham '
  'gia của các em hoàn toàn tự nguyện và các em có thể tự do rút lui khỏi việc '
  'tham gia bất cứ lúc nào mà không có hậu quả.')

P('Một lo ngại khác là cho phúc lợi của người tham gia học sinh. Mặc dù khó '
  'xảy ra, có thể học sinh có thể trở nên buồn bực khi trả lời khảo sát hoặc '
  'khi tham gia một nhóm tập trung. Nhóm thu thập dữ liệu bao gồm 2 nhà tâm lý '
  'học cấp thạc sĩ và được dẫn dắt bởi một nhà tâm lý học lâm sàng tiến sĩ '
  '(Đặng Hoàng Minh). Nhóm thu thập dữ liệu đã được đào tạo để tìm kiếm các '
  'dấu hiệu đau khổ của học sinh và để đưa ra hỗ trợ hoặc giới thiệu nếu cần. '
  'Không có sự cố đau khổ nào được báo cáo.')

P('Giao thức nghiên cứu đã nhận được rà soát và phê duyệt đạo đức từ Hội đồng '
  'Đạo đức Nghiên cứu (IRB) của Đại học Quốc gia Việt Nam. Rà soát IRB bao gồm '
  'các mẫu đồng ý tham gia cho học sinh và phụ huynh.')

doc.add_page_break()

# ============================================================
# CHUONG 5 — PHAT HIEN NGHIEN CUU
# ============================================================
H('CHƯƠNG 5: PHÁT HIỆN NGHIÊN CỨU VỀ SỨC KHOẺ TÂM THẦN VÀ HẠNH PHÚC HỌC SINH VỊ '
  'THÀNH NIÊN VIỆT NAM', level=1)
P('(Chapter 5 — Adolescent Student Mental Health and Well-Being in Viet Nam Study Findings)',
  italic=True, align='center', size=11)

# --- Trang 36 ---
PageMark('--- Trang 36, UNICEF Việt Nam, 2022 ---')

H('Các phát hiện chính (Key Findings)', level=2)

P('1. Các triệu chứng sức khoẻ tâm thần tự báo cáo của học sinh chỉ ra rằng '
  'khoảng 26 % học sinh vị thành niên có nguy cơ trung bình hoặc cao về các vấn '
  'đề sức khoẻ tâm thần. Các vấn đề với bạn bè (bao gồm trải nghiệm bắt nạt) và '
  'các vấn đề cảm xúc (các triệu chứng trầm cảm và lo âu) là các vấn đề phổ biến '
  'nhất ở mức 32 % và 31 % tương ứng. Khoảng 14 % học sinh báo cáo các triệu '
  'chứng tăng động, bao gồm bốc đồng và tập trung kém, và 11 % học sinh báo cáo '
  'các vấn đề hành vi, bao gồm bất tuân và nói dối.', bold=True)

P('2. Các em gái vị thành niên báo cáo tỷ lệ các vấn đề cảm xúc cao hơn đáng kể '
  '(các triệu chứng trầm cảm và lo âu) so với các em trai.', bold=True)

P('3. Phụ huynh thường thiếu kiến thức về các vấn đề sức khoẻ tâm thần phổ biến '
  'và sự phát triển bình thường của trẻ em, và mong muốn có thông tin.', bold=True)

P('4. Tất cả các hiệu trưởng, các nhà quản lý trường học và các chuyên gia cấp '
  'Bộ (từ cả 3 ngành) được phỏng vấn đều bày tỏ quan ngại nghiêm trọng về sức '
  'khoẻ tâm thần và hạnh phúc của học sinh vị thành niên, cùng với mối quan ngại '
  'về tác động tiêu cực của các vấn đề sức khoẻ tâm thần học sinh lên việc học '
  'tập và chức năng xã hội – cảm xúc của học sinh.', bold=True)

P('5. Các em gái thường được xem là dễ bị tổn thương hơn với các vấn đề tâm lý, '
  'và hành vi của người lớn đối với các em gái có thể "bảo vệ" các em theo cách '
  'không phải lúc nào cũng hỗ trợ khả năng phục hồi và sức mạnh MHPSS của nữ '
  'giới.', bold=True)

P('6. Giáo viên và các nhà quản lý lo ngại về sự cô lập xã hội của học sinh dân '
  'tộc thiểu số, tỷ lệ tìm kiếm trợ giúp thấp hơn và nguy cơ tự tử.', bold=True)

P('7. Học sinh LGBTQ lo ngại về sức khoẻ tâm thần và hạnh phúc của mình. Các mối '
  'quan hệ gia đình và nỗi sợ bị kỳ thị và phân biệt đối xử là các nguồn gây '
  'stress và thách thức sức khoẻ tâm thần.', bold=True)

# --- Trang 37 ---
PageMark('--- Trang 37, UNICEF Việt Nam, 2022 ---')

H('5.1. Cuộc sống học sinh (Student Life)', level=2)

P('Học sinh là những người tham gia nhiệt tình và có nhiều điều để chia sẻ về '
  'cuộc sống, sức khoẻ tâm thần và hạnh phúc của mình. Học sinh vị thành niên '
  'thường đi học 6 ngày/tuần (bao gồm nửa ngày Thứ Bảy). Hầu hết học sinh báo '
  'cáo tham gia hai lớp học thêm trở lên ngoài trường mỗi tuần, bao gồm các lớp '
  'toán, lý và tiếng Anh. Khi vị thành niên không ở trường hoặc trong các lớp '
  'học thêm, các em tham gia nhiều hoạt động khác nhau, bao gồm học tập, việc '
  'nhà và các hoạt động giải trí.')

P('Khi được hỏi các em làm gì để thư giãn, học sinh thường nói rằng các em dành '
  'thời gian trên mạng xã hội và nhắn tin với bạn bè qua các ứng dụng nhắn tin. '
  'Những học sinh khác chia sẻ rằng các em lướt Facebook hoặc "hóng tin" (theo '
  'dõi tin tức giải trí). Một số học sinh chơi thể thao, bao gồm cầu lông, bóng '
  'đá, bóng rổ, yoga và bơi lội. Một vài học sinh báo cáo chơi nhạc cụ, hát và '
  'sáng tác nhạc. Những học sinh khác thích nghe nhạc, xem phim, chụp ảnh và '
  'đọc sách. Học sinh coi trọng thời gian gia đình như được thể hiện qua một '
  'học sinh chia sẻ cảm thấy "rất hạnh phúc" khi cả gia đình cùng chơi cầu lông. '
  'Một số học sinh lưu ý tác động tích cực của các hoạt động lên sức khoẻ và '
  'hạnh phúc của mình. Như một học sinh ở Hà Nội chia sẻ: "Em không có nhiều '
  'thời gian, nhưng giữa các lớp học có hai giờ để em nghỉ ngơi. Em sử dụng thời '
  'gian đó để chơi với chó và tập thể dục. Sau đó em nhận thấy cơ thể khoẻ hơn '
  'và cảm thấy thư giãn hơn một chút."')

P('Ngoài các trách nhiệm học tập, học sinh báo cáo nhiều trách nhiệm ở nhà. Học '
  'sinh giúp việc nhà, bao gồm "giặt giũ, quét nhà và dọn bàn." Một vài học sinh '
  'nấu bữa tối, và nhiều học sinh chịu trách nhiệm trông em nhỏ. Một học sinh '
  'từ Điện Biên báo cáo: "[Em] không có thời gian đi chơi. Cha mẹ yêu cầu em '
  'trông em gái đến khi em ngủ. Bởi vì nếu em đi ra ngoài, em sẽ không biết khi '
  'nào em bé khóc." Trong khi nhiều học sinh cảm thấy mức độ công việc nhà là '
  'hợp lý, một vài học sinh cho biết các em cảm thấy choáng ngợp bởi những '
  'trách nhiệm này. Một học sinh chia sẻ: "Làm chị cả là điều em thấy khó khăn '
  'vì bố mẹ em có quá nhiều kỳ vọng và đôi khi bỏ bê em."')

P('Học sinh chia sẻ những thách thức tại nhà bao gồm stress liên quan đến xung '
  'đột gia đình. Một số học sinh báo cáo rằng cha mẹ thường tức giận, thường '
  '"la hét" và "chửi thề." Một học sinh mô tả: "khi cha mẹ em về nhà từ công '
  'việc, họ mệt mỏi và họ trút giận lên em. Họ tức giận nếu em không nghe lời. '
  'Họ dễ giận."')

H('5.2. Sức khoẻ tâm thần và hạnh phúc học sinh '
  '(Student Mental Health and Well-being)', level=2)

H('5.2.1. Nhận thức của học sinh về sức khoẻ tâm thần và hạnh phúc của mình '
  '(Student Perceptions of their Mental Health and Well-Being)', level=3)

P('Trong các thảo luận về sức khoẻ tâm thần, học sinh thường xác định cảm giác '
  'bị stress hoặc mệt mỏi. Khi được yêu cầu chia sẻ thêm về tác động của stress, '
  'học sinh thảo luận về cảm giác sợ hãi, lo âu, thất vọng và buồn bã, chủ yếu '
  'liên quan đến thành tích học tập và tương lai của các em. Các em lưu ý rằng '
  'những cảm giác này có thể làm suy giảm khả năng học tập và tương tác với '
  'bạn bè của mình.')

P('Ngoài "stress," học sinh không thể hiện nhiều kiến thức về sức khoẻ tâm '
  'thần. Khi được hỏi về trầm cảm, học sinh ở Điện Biên trả lời rằng đó là '
  '"sợ mọi thứ xung quanh," "có quá nhiều áp lực rồi các em trở nên im lặng và '
  'không muốn liên lạc hoặc nói chuyện với bất kỳ ai," "tức giận bất cứ khi nào '
  'nhìn thấy ai đó," và "xa cách." Và học sinh không biết về bản chất và các '
  'triệu chứng của các rối loạn lo âu.')

P('Học sinh bày tỏ lo ngại về sự kỳ thị và phân biệt đối xử đối với sức khoẻ '
  'tâm thần. Khi được hỏi giáo viên, phụ huynh và các người lớn khác nhìn nhận '
  'những người có vấn đề sức khoẻ tâm thần như thế nào, học sinh từ Điện Biên '
  'nói "Họ bàn tán. Họ nói những người đó có vấn đề gì" và "họ phân biệt đối '
  'xử." Khi được hỏi liệu giáo viên có giúp đỡ một học sinh đang có vấn đề cảm '
  'xúc không, những người tham gia nói có, các thầy cô sẽ "động viên và an ủi."')

P('Vị thành niên chia sẻ các chiến lược mà các em biết hoặc sử dụng để quản lý '
  'stress và cải thiện sức khoẻ tâm thần của mình. Học sinh thường nói nhất '
  'rằng các em sẽ nói chuyện hoặc nhắn tin với một người bạn khi cảm thấy buồn '
  'hoặc stress.')

# --- Trang 38 ---
PageMark('--- Trang 38, UNICEF Việt Nam, 2022 ---')
P('Với một học sinh lưu ý: "Bạn bè của em có thể giúp em giải toả stress và nỗi '
  'buồn." Một số thanh thiếu niên chia sẻ rằng các em thấy nói chuyện với gia '
  'đình và giáo viên là hữu ích. Học sinh biết rằng các em có thể nhận được các '
  'loại hỗ trợ khác nhau từ những người khác nhau, như một học sinh ở Gia Lai '
  'giải thích: "Cô chủ nhiệm của em sẽ giải thích cho em cách nhìn mọi thứ một '
  'cách tích cực, trong khi bạn thân của em sẽ nói thẳng với em theo một cách '
  'tàn nhẫn để khiến nó rất tiêu cực để em đi làm gì đó về nó." Một số học sinh '
  'cũng nêu rằng các em tìm đến thiên nhiên để giảm stress, nói rằng các em '
  'thích được ở trong vườn hoặc đi bộ trong không khí trong lành. Hai học sinh '
  'từ Điện Biên chia sẻ rằng các em thấy hữu ích khi viết cảm xúc vào nhật ký '
  'khi bị stress hoặc buồn bã. Các học sinh khác chia sẻ rằng khi các em chơi '
  'video game, "nỗi buồn biến mất." Và học sinh hiểu rằng sức khoẻ thể chất '
  'quan trọng đối với sức khoẻ tâm thần với một học sinh từ Đồng Tháp chia sẻ: '
  '"bạn có thể ăn đồ ăn ngon, với nhiều vitamin để giúp thắp sáng tinh thần."')

P('Dữ liệu định lượng cung cấp thêm thông tin về sức khoẻ tâm thần và hạnh phúc '
  'của học sinh vị thành niên (xem Phụ lục 3 để có bảng và kết quả Dữ liệu Định '
  'lượng đầy đủ). Các phát hiện từ SDQ25 đánh giá sức khoẻ tâm thần học sinh '
  'chỉ ra rằng 26,1 % học sinh vị thành niên có nguy cơ trung bình hoặc cao về '
  'các vấn đề sức khoẻ tâm thần. Các loại vấn đề phổ biến nhất được báo cáo là '
  'các vấn đề bạn bè (ví dụ: "Những đứa trẻ hoặc người trẻ khác chọc ghẹo hoặc '
  'bắt nạt tôi," "Tôi hoà hợp tốt với người lớn hơn là với người cùng tuổi"), '
  'được tìm thấy ở 32 % học sinh, và các vấn đề cảm xúc (ví dụ: "Nhiều lo lắng '
  'hoặc thường xuyên có vẻ lo lắng," "Thường xuyên không vui, chán nản hoặc '
  'khóc"), được tìm thấy ở 30,9 % học sinh. Các triệu chứng tăng động (ví dụ: '
  '"dễ bị phân tâm, sự tập trung bị lung lay," "Suy nghĩ mọi thứ trước khi hành '
  'động") được báo cáo bởi 14,4 % học sinh và các vấn đề hành vi (ví dụ: "Nhìn '
  'chung vâng lời, thường làm những gì người lớn yêu cầu," "Thường xuyên nói '
  'dối hoặc lừa gạt") được tìm thấy ở 11 % học sinh. Các em trai và em gái báo '
  'cáo tỷ lệ vấn đề tương tự, ngoại trừ Vấn đề Cảm xúc (các triệu chứng trầm '
  'cảm và lo âu), nơi chúng ta thấy tỷ lệ triệu chứng vấn đề cảm xúc cao hơn '
  'đáng kể ở các em gái so với các em trai. Cấp học cũng là yếu tố chính trong '
  'sức khoẻ tâm thần học sinh vị thành niên với học sinh trung học phổ thông '
  'cho thấy nhiều vấn đề hành vi, tăng động, cảm xúc và tổng vấn đề đáng kể hơn '
  'so với học sinh trung học cơ sở.')

H('5.2.2. Nhận thức của giáo viên về sức khoẻ tâm thần và hạnh phúc học sinh '
  '(Teacher Perceptions of Student Mental Health and Well-Being)', level=3)

P('Giáo viên nhìn chung đồng ý rằng hầu hết học sinh nhìn chung đều khoẻ mạnh và '
  'hạnh phúc. Như một giáo viên từ Hà Nội báo cáo: "Hầu hết thời gian, giáo viên '
  'thấy học sinh ở độ tuổi năng lượng cao với triển vọng lạc quan... Tôi nghĩ '
  'môi trường học tập đã khuyến khích học sinh cảm thấy thoải mái và vui vẻ."')

P('Tuy nhiên, giáo viên có bày tỏ một số lo ngại về sức khoẻ tâm thần học sinh. '
  'Như một giáo viên từ Đồng Tháp nêu: "Tôi nghĩ rằng khoảng 85 % học sinh có vẻ '
  'vui vẻ. Tuy nhiên, khoảng 10 % hoặc 15 % bị mắc kẹt trong tình huống ngược '
  'lại. Các em luôn đeo khẩu trang cho bất kỳ dịp nào, như thể các em đang chịu '
  'đựng trầm cảm. Các em cũng không thể tập trung lắng nghe những gì chúng tôi '
  'dạy trong lớp. Các em thường nhìn quanh lớp học và giữ im lặng thay vì nói '
  'chuyện hoặc chơi với bạn bè, và khi chúng tôi đến gần để nói chuyện với các '
  'em, các em giật mình." Một giáo viên ở Hà Nội đồng ý với mối quan ngại dành '
  'cho học sinh "hướng nội" ít khi nói chuyện hoặc bộc lộ cảm xúc. Các giáo '
  'viên khác chia sẻ mối quan ngại dành cho học sinh thiếu tự tin và thiếu gắn '
  'kết trong lớp.')

P('Giáo viên xác định các khiếm khuyết nền tảng của học sinh mà họ tin là dẫn '
  'đến các vấn đề sức khoẻ tâm thần. Một giáo viên ở Hà Nội bày tỏ lo ngại về '
  'khả năng giải quyết vấn đề của học sinh và thấy rằng sự thiếu hụt này dẫn '
  'đến stress gia tăng: "học sinh ngày nay không quá sắc bén hoặc thuần thục '
  'trong việc tìm các giải pháp để vượt qua hoặc sửa chữa các hạn chế của '
  'mình." Một giáo viên khác đồng ý rằng học sinh thường thiếu khả năng giải '
  'quyết vấn đề và động lực nội tại cho việc học, và tin rằng các vấn đề này '
  'dẫn đến stress và các vấn đề sức khoẻ tâm thần cản trở việc học.')

P('Giáo viên đôi khi chia sẻ sự bối rối về bản chất các vấn đề học sinh. Một '
  'giáo viên từ Đồng Tháp chia sẻ ví dụ về một học sinh đang nói chuyện một mình '
  'ở nhà và không gắn kết với trường học và tương tác xã hội. Cô đã lo ngại và '
  'nghi ngờ vấn đề là trầm cảm.')

# --- Trang 39 ---
PageMark('--- Trang 39, UNICEF Việt Nam, 2022 ---')
P('Cô đã không giới thiệu học sinh đến bệnh viện nhưng đã dành nhiều sự chú ý '
  'hơn cho học sinh đó ở trường.')

P('Dữ liệu định lượng giáo viên từ 66 giáo viên cung cấp thêm cái nhìn sâu hơn '
  'vào nhận thức của giáo viên về sức khoẻ tâm thần và hạnh phúc học sinh vị '
  'thành niên. Kết quả khảo sát chỉ ra rằng hầu hết tất cả giáo viên đều lo '
  'ngại về stress và hạnh phúc của học sinh (91 %), với 53 % giáo viên báo cáo '
  'họ "Rất lo ngại." Tương tự, hầu hết tất cả giáo viên đều lo ngại về sức '
  'khoẻ tâm thần của học sinh (95 %), với 54 % giáo viên báo cáo họ "Rất lo '
  'ngại." Hầu hết giáo viên (34 giáo viên) ước tính rằng 10 % hoặc ít hơn học '
  'sinh của họ đã trải qua các vấn đề stress trong tháng qua, mặc dù 15 giáo '
  'viên ước tính tỷ lệ cao hơn từ 10–20 % học sinh có vấn đề stress trong tháng '
  'qua. Giáo viên cũng báo cáo tỷ lệ thấp các vấn đề sức khoẻ tâm thần ở học '
  'sinh, với hầu hết giáo viên (50 giáo viên) ước tính 10 % hoặc ít hơn học '
  'sinh có vấn đề sức khoẻ tâm thần trong tháng qua. Các vấn đề sức khoẻ tâm '
  'thần học sinh vị thành niên phổ biến nhất được giáo viên báo cáo bao gồm '
  'Chú ý Kém, Thiếu Tự Tin, Năng lượng Thấp/Mệt mỏi và Lo âu. Các vấn đề phổ '
  'biến vừa phải được giáo viên báo cáo bao gồm Tức giận/Hung hăng, Buồn/Trầm '
  'cảm và Hành vi Gây rối. Giáo viên chỉ ra rằng các vấn đề ít phổ biến nhất '
  'là Hành vi Đối lập hoặc Thách thức, Lạm dụng Rượu/Chất và Xung đột Bạn bè '
  'Nghiêm trọng.')

H('5.2.3. Nhận thức của phụ huynh về sức khoẻ tâm thần và hạnh phúc học sinh '
  '(Parent Perceptions of Student Mental Health and Well-Being)', level=3)

P('Nhìn chung, hầu hết phụ huynh báo cáo hài lòng với hạnh phúc của con mình, '
  'thấy các em thường vui vẻ và ngoan ngoãn. Khi phụ huynh được hỏi về những lo '
  'ngại sức khoẻ tâm thần học sinh, phụ huynh thường ban đầu nói về các vấn đề '
  'liên quan đến dậy thì, các mối quan hệ tình dục/lãng mạn, sử dụng điện thoại '
  'và sự không hài lòng của trẻ về việc học. Nhiều phụ huynh lo ngại về sự cáu '
  'kỉnh, tức giận hoặc thách thức của con ở nhà. Phụ huynh bày tỏ lo ngại về '
  'việc trẻ không tôn trọng. Như một phụ huynh từ Đồng Tháp chia sẻ: "Có cái '
  'tính bộc phát này, so với khi tôi ở cùng tuổi, có [nhiều] bộc phát và bất '
  'đồng hơn. Nhìn chung, tôi thấy cháu khá là bướng bỉnh." Một số phụ huynh '
  'tin rằng thiếu tôn trọng cha mẹ là vấn đề hiện đại và lo ngại rằng điều đó '
  'phản ánh giá trị đang thay đổi. Như một phụ huynh lưu ý: "Trẻ con bây giờ '
  'dễ tức giận [so với] chúng tôi khi chúng tôi ở tuổi đó. Ví dụ, ngày xưa, '
  'chúng tôi sẽ im lặng và phớt lờ khi cha mẹ quát mắng chúng tôi, nhưng bây '
  'giờ chỉ một hai câu là trẻ cãi lại." Phụ huynh cũng báo cáo thiếu tự tin '
  'trong cách phản ứng với sự thiếu tôn trọng và tức giận của trẻ em.')

P('Ban đầu, phụ huynh thường không nghĩ đến các triệu chứng sức khoẻ tâm thần '
  'phổ biến như mất năng lượng, thiếu hứng thú, lo lắng, v.v. Khi phụ huynh ở '
  'Điện Biên được hỏi liệu con trung học cơ sở của họ có trải qua stress '
  'không, ban đầu họ nghĩ là không. Tuy nhiên, khi người dẫn dắt mô tả các '
  'dấu hiệu và triệu chứng của stress, các phụ huynh sau đó báo cáo rằng con '
  'họ thực sự trải qua stress. Các mô hình phản hồi này chỉ ra rằng phụ huynh '
  'thiếu hiểu biết về các vấn đề sức khoẻ tâm thần phổ biến và các triệu chứng '
  'liên quan.')

P('Phụ huynh bày tỏ sự bối rối về sự phát triển bình thường của vị thành niên '
  'và không chắc chắn khi nào các hành vi chỉ ra vấn đề. Phụ huynh ở Gia Lai '
  'báo cáo các vấn đề như cáu kỉnh, phân tâm và mệt mỏi, và tự hỏi liệu những '
  'hành vi này có thể do dậy thì không. Một phụ huynh từ Điện Biên hỏi người '
  'dẫn dắt phải làm gì về sự thiếu tôn trọng của vị thành niên đối với cha mẹ. '
  'Phụ huynh từ mỗi tỉnh báo cáo các dấu hiệu lo âu và trầm cảm nhưng không '
  'chắc chắn làm thế nào để hỗ trợ con mình tốt nhất hoặc cách tìm kiếm thông '
  'tin hoặc hỗ trợ.')

H('5.2.4. Nhận thức của người cung cấp thông tin chủ chốt về sức khoẻ tâm thần '
  'và hạnh phúc học sinh (Key Informant Perceptions)', level=3)

P('Tất cả các hiệu trưởng được phỏng vấn đều bày tỏ lo ngại về sức khoẻ tâm '
  'thần và hạnh phúc học sinh vị thành niên. Thảo luận về các vấn đề sức khoẻ '
  'tâm thần phổ biến ở học sinh bao gồm stress, lo âu và trầm cảm, nhưng cũng '
  'có một số thảo luận về các hành vi được coi là vô đạo đức. Ví dụ, một hiệu '
  'trưởng từ Gia Lai báo cáo học sinh hút thuốc, có các mối quan hệ không phù '
  'hợp dẫn đến xung đột, và dành thời gian trên điện thoại hoặc internet. Một '
  'hiệu trưởng khác lo ngại về việc học sinh thiếu động lực. Bà và nhân viên '
  'nhà trường vật lộn với cách tạo động lực cho học sinh và thấy rằng can '
  'thiệp gia đình không giúp nhiều.')

# --- Trang 40 ---
PageMark('--- Trang 40, UNICEF Việt Nam, 2022 ---')
AddImg('p040_img1_Im0.jpg',
       'Hình 4 (trang 40). Học sinh nữ trong sân trường — chương Phát hiện',
       'Schoolgirl in schoolyard — Findings chapter')
P('Một số hiệu trưởng chia sẻ lo ngại về các khiếm khuyết kỹ năng mà họ tin là '
  'nằm sau các vấn đề sức khoẻ tâm thần và chức năng của học sinh. Một hiệu '
  'trưởng từ Điện Biên báo cáo rằng học sinh thường thiếu các kỹ năng giải '
  'quyết vấn đề cần thiết và nêu: "Nếu học sinh không được trang bị kỹ năng '
  'giải quyết vấn đề, các em chắc chắn sẽ đối mặt với các tình huống mà các '
  'em không thể xử lý các vấn đề trong tình bạn và các mối quan hệ với giáo '
  'viên."')

P('Các hiệu trưởng cũng chia sẻ lo ngại về tác động tiêu cực của các vấn đề '
  'sức khoẻ tâm thần học sinh lên việc học và chức năng xã hội – cảm xúc. Như '
  'một hiệu trưởng trung học phổ thông từ Hà Nội báo cáo: "[các vấn đề sức '
  'khoẻ tâm thần] giảm khả năng học tập, khả năng tập trung và xử lý tình '
  'huống của các em. Các em sẽ không quan tâm đến bất kỳ môn học nào hoặc bất '
  'kỳ lựa chọn nghề nghiệp nào trong tương lai." Một hiệu trưởng khác mô tả '
  'vòng luẩn quẩn các vấn đề sức khoẻ tâm thần và thành tích kém: "học sinh '
  'trở nên cáu kỉnh, nóng nảy và phản ứng quá mức. [Khi các em đối mặt với] '
  'những điều nhẹ nhàng và tình cờ, bây giờ các em cảm thấy như đang bị đâm, '
  'bị gọi tên, bị chỉ trích. Rồi các em không tập trung, kết quả của các em '
  'giảm xuống, và điều này lặp đi lặp lại. Các em hoàn toàn mất sự tự tin mà '
  'các em đã có trước đây, và khả năng trước đây của các em không còn nữa."')

P('Các hiệu trưởng báo cáo rằng các vấn đề sức khoẻ tâm thần thường bị hiểu '
  'lầm và xác định sai ở trường. Một hiệu trưởng từ Hà Nội chia sẻ ví dụ về '
  'một học sinh lớp 8 có vấn đề ngủ nghiêm trọng và rút khỏi trường học. '
  'Trong ví dụ này, nhà trường và phụ huynh đã bỏ lỡ dấu hiệu của bệnh tâm '
  'thần mới nổi nên không tìm kiếm sự trợ giúp chuyên nghiệp. Đứa trẻ đã phát '
  'triển bệnh tâm thần nghiêm trọng, sau đó phải nhập viện và không bao giờ '
  'trở lại trường.')

P('Tất cả các nhà quản lý Sở GD&ĐT được phỏng vấn đều chia sẻ những mối quan '
  'ngại này về sức khoẻ tâm thần học sinh vị thành niên. Một nhà quản lý từ '
  'Hà Nội bày tỏ tầm quan trọng của việc hỗ trợ sức khoẻ tâm thần học sinh vị '
  'thành niên: "Tôi nghĩ sức khoẻ tâm thần nên được coi là quan trọng hơn và '
  'không thể tranh cãi đối với sự phát triển của học sinh, có thể thậm chí '
  'được xếp hạng ở tầm quan trọng cao hơn so với thành tích học tập. Tôi '
  'nghĩ tốt hơn là dạy học sinh cách kiểm soát bản thân trước khi ném cho các '
  'em kiến thức toán học như sin, cos, tan, cotan."')

P('Các nhà quản lý Sở Y tế cũng xem sức khoẻ tâm thần vị thành niên là một '
  'vấn đề nghiêm trọng.')

# --- Trang 41 ---
PageMark('--- Trang 41, UNICEF Việt Nam, 2022 ---')
P('Các chuyên gia chia sẻ rằng trầm cảm và lo âu là các vấn đề ở học sinh. '
  'Một người tham gia từ Đồng Tháp nêu: "Có những stress mà các em giữ bên '
  'trong và không nói ra vì các em thấy quá nhiều kỳ vọng từ nhà trường và '
  'cha mẹ, trong khi bản thân các em không thể đáp ứng những kỳ vọng này. '
  'Đôi khi điều này dẫn đến trạng thái ẩn mình, rút lui khỏi mọi thứ nên '
  'cuối cùng các em không giao tiếp." Các quan chức y tế thường thảo luận '
  'về tác động của các vấn đề sức khoẻ tâm thần lên sự phát triển vị thành '
  'niên. Một quan chức Sở Y tế từ Điện Biên nhận xét: "Vấn đề sức khoẻ tâm '
  'lý – tâm thần ở tuổi vị thành niên thực sự rất quan trọng vì giai đoạn '
  'này có tác động lớn đến các mục tiêu giáo dục của vị thành niên. Sức '
  'khoẻ tâm thần quyết định liệu các em có bị phân tâm bởi điều gì đó '
  'không hoặc liệu các em có đạt được định hướng cần thiết để thành công '
  'trong cuộc sống hay không. Một khi vị thành niên ổn định về mặt tâm lý, '
  'các em có thể được định hướng vào con đường nghề nghiệp đúng. Nếu các '
  'em không có hướng dẫn để kiểm soát các phân tâm tâm lý, cuộc sống của '
  'các em có thể đi lạc lối."')

P('Tương tự, các nhà quản lý Sở LĐ-TB&XH bày tỏ lo ngại về sức khoẻ tâm thần '
  'học sinh vị thành niên. Một chuyên gia Sở LĐ-TB&XH Hà Nội thấy rằng vị '
  'thành niên thiếu khả năng giải quyết vấn đề và phát triển các mối quan '
  'hệ, và thấy những khiếm khuyết kỹ năng này đặt học sinh vào nguy cơ: '
  '"Vấn đề đầu tiên là trẻ em khá yếu về khả năng giải quyết vấn đề trong '
  'cuộc sống và khả năng kết nối với môi trường xung quanh. Tôi nghĩ các '
  'mối quan hệ và kết nối xã hội, và tìm kiếm trợ giúp của các em không tốt '
  'và không đủ người chú ý. [Những vấn đề này] giải thích tại sao có những '
  'vụ tự tử làm cha mẹ và bạn bè của đứa trẻ ngạc nhiên, vì nhiều người '
  'trong số họ cảm thấy đứa trẻ chỉ đang bị stress."')

P('Các chuyên gia từ Bộ Y tế, Bộ GD&ĐT và Bộ LĐ-TB&XH đồng ý rằng sức khoẻ '
  'tâm thần học sinh vị thành niên đại diện cho một thách thức nghiêm trọng '
  'đối với xã hội. Một chuyên gia từ Bộ Y tế báo cáo sự gia tăng các vấn đề '
  'sức khoẻ tâm thần trong 10–20 năm qua do "phát triển kinh tế – xã hội, '
  'công nghiệp hoá, hiện đại, phát triển công nghệ thông tin và thay đổi '
  'trong tiêu chuẩn sống. Những thay đổi này đặt áp lực lên thế hệ trẻ, '
  'thường dẫn đến các vấn đề." Chuyên gia này cũng lưu ý rằng các vấn đề '
  'sức khoẻ tâm thần có vẻ phổ biến hơn bây giờ do nhận thức tăng lên về '
  'các vấn đề này. Ông nêu: "Trong quá khứ, có thể đã có một vài trường '
  'hợp [bệnh tâm thần vị thành niên] nhưng mọi người không quan tâm nên họ '
  'không nhận ra vấn đề. Bây giờ kiến thức của mọi người đang mở rộng, và '
  'họ có nhiều khả năng xác định các trường hợp đó hơn."')

H('5.3. Dữ liệu và nhận thức về cách nhu cầu sức khoẻ tâm thần học sinh khác '
  'biệt theo giới tính, dân tộc và tình trạng LGBTQ '
  '(Data and perceptions regarding how student mental health needs vary by '
  'gender, ethnicity, and LGBTQ status)', level=2)

P('Niềm tin về sự khác biệt giới tính liên quan đến sức khoẻ tâm thần vị '
  'thành niên đôi khi được bày tỏ trong các FGD và KII. Nhìn chung, phụ '
  'huynh và các nhà quản lý chia sẻ nhận thức về các em gái là dễ bị tổn '
  'thương hơn với các vấn đề sức khoẻ tâm thần. Một phụ huynh ở Điện Biên '
  'bày tỏ niềm tin rằng các em gái không khoẻ mạnh như các em trai, nêu: '
  '"Tôi có một cô con gái và bạn biết con gái không khoẻ bằng những người '
  'khác, đó là lý do tại sao các kỳ thi luôn làm tôi lo lắng. Hơn nữa, tôi '
  'có giúp con gái mọi thứ, những thứ con có thể không làm được mọi lúc, '
  'chẳng hạn như giúp con với việc nhà để thời gian còn lại có thể dành '
  'riêng cho việc học. Điều này là do tình trạng sức khoẻ của con gái tôi. '
  'Nếu không, con có thể có dấu hiệu mất ngủ phải được hỗ trợ bằng thuốc." '
  'Một phụ huynh khác từ Gia Lai bày tỏ niềm tin rằng các em trai ít bị '
  'ảnh hưởng bởi các vấn đề sức khoẻ tâm thần, suy đoán rằng đó là vì "các '
  'em trai thường dạn dĩ hơn."')

P('Sức khoẻ tâm thần và hạnh phúc của học sinh dân tộc thiểu số là mối quan '
  'ngại của một số người tham gia. Mối quan ngại đặc biệt là sự cô lập xã '
  'hội của học sinh dân tộc thiểu số, thiếu tìm kiếm trợ giúp và nguy cơ tự '
  'tử. Một nhà quản lý Sở GD&ĐT ở Điện Biên lưu ý rằng học sinh dân tộc '
  'thiểu số trong huyện của mình dễ bị tổn thương với các vấn đề sức khoẻ '
  'tâm thần: "Học sinh Mông chỉ xem việc đi chơi với học sinh cùng dân tộc, '
  'và kết quả là các em hiếm khi chơi hoặc làm quen với người Thái. Do đó, '
  'các vấn đề vẫn nằm trong một nhóm dân tộc cụ thể, đó là lý do tại sao '
  'rất khó để giải quyết mọi thứ. Học sinh hiếm khi chủ động chia sẻ lo '
  'lắng của mình và sẽ thích tự xử lý sự kiện một mình. Và nhiều lần, cách '
  'các em đối phó với những sự kiện này có tính chất cực kỳ tiêu cực.')

# --- Trang 42 ---
PageMark('--- Trang 42, UNICEF Việt Nam, 2022 ---')
P('Ví dụ, trong trường hợp các vấn đề tình yêu lãng mạn, một cô gái hoặc một '
  'chàng trai có thể có xung đột trong mối quan hệ và quyết định tự cô lập '
  'và không chia sẻ câu chuyện với bất kỳ ai. Điều này có thể dẫn đến ý '
  'tưởng tự tử và tử vong. Học sinh có một nguồn lực tự nhiên là lá ngón '
  '(cỏ đau tim có độc), có thể phục vụ cho ý tưởng tự tử khá dễ dàng. Để '
  'kết luận, điều khó khăn trong việc hiểu được sức khoẻ tâm thần của một '
  'học sinh nằm ở chính các học sinh, những người không có kỹ năng chia sẻ '
  'về các vấn đề cá nhân khó khăn để được giúp tháo gỡ tâm trí của mình." '
  'Một nhà quản lý Sở Y tế ở Điện Biên cũng bày tỏ lo ngại về tự tử ở vị '
  'thành niên Mông qua ngộ độc "lá ngón." Sự miễn cưỡng của gia đình trong '
  'việc tìm kiếm trợ giúp cho các vấn đề của trẻ em cũng được lưu ý như '
  'một lo ngại về học sinh dân tộc thiểu số. Một nhà quản lý Sở LĐ-TB&XH '
  'từ Điện Biên lưu ý rằng các cộng đồng dân tộc thiểu số ít có khả năng '
  'tiếp cận chăm sóc sức khoẻ khi con cái bị bệnh, dẫn đến các trường hợp '
  'bệnh nghiêm trọng (bao gồm động kinh được quản lý kém) ở học sinh vị '
  'thành niên đôi khi.')

P('Học sinh LGBTQ báo cáo lo ngại về sức khoẻ tâm thần và hạnh phúc của '
  'mình. Mặc dù nghiên cứu chỉ bao gồm một nhóm nhỏ học sinh LGBTQ, kết '
  'quả chỉ ra nhu cầu lo ngại và nghiên cứu thêm. Học sinh chia sẻ cảm '
  'giác lo âu về các mối quan hệ và thành tích học tập, hình ảnh bản thân '
  'tiêu cực và bất an, và sự bất ổn cảm xúc. Học sinh báo cáo cảm thấy '
  'choáng ngợp bởi các cảm xúc tiêu cực đôi khi. Ví dụ, một học sinh chia '
  'sẻ: "Em thường đi theo hướng tiêu cực hơn khi quá stress, và em không '
  'biết cách giải toả, không có cách nào em có thể quên, nên sau đó em tự '
  'hại một chút. Trước đây, em thường dùng dao. Rồi sau đó, em không muốn '
  'cắt vào cơ thể mình nên em vứt mọi thứ sắc trong nhà đi, rồi em bắt '
  'đầu chuyển sang các phương pháp khác — như đập đầu vào tường hoặc đổ '
  'nước sôi lên tay." Trong khi học sinh này báo cáo đã nhận được một số '
  'tư vấn tâm lý, các em không thấy nó hữu ích vì các em nhận thấy tư vấn '
  'viên phán xét và không đồng cảm.')

P('Các mối quan hệ gia đình có thể là một nguồn stress đáng kể cho vị '
  'thành niên LGBTQ. Học sinh báo cáo lo âu về việc công khai với cha mẹ. '
  'Một số học sinh chia sẻ rằng phản ứng tiêu cực của cha mẹ đối với xu '
  'hướng tính dục và bản dạng giới của các em là nguồn gốc của buồn bã '
  'và đau khổ. Như một học sinh bày tỏ: "Đối với em, [mối quan hệ] gây '
  'stress nhất là mối quan hệ với gia đình. Thực ra, em đã độc lập với '
  'gia đình khá sớm. Từ trung học cơ sở khi em chọn ở xa gia đình đến '
  'bây giờ, em không nói chuyện nhiều với mẹ, bố hay anh chị em. Khi em '
  'công khai, bố em và gia đình cắt mọi hỗ trợ cho em, bao gồm cả hỗ trợ '
  'tài chính cho việc học của em. Họ có phần chấp nhận nhưng vẫn chưa '
  'ủng hộ em như một người đồng tính." Học sinh có cha mẹ cuối cùng ủng '
  'hộ và chấp nhận các em là LGBTQ chia sẻ rằng đây là nguồn gốc của khả '
  'năng phục hồi.')

P('Trong khi công khai có thể là một quá trình gây stress cho vị thành '
  'niên, một số vị thành niên thấy nó cuối cùng là tích cực. Một học sinh '
  'chia sẻ: "Cá nhân em cảm thấy công khai chính thức là một người LGBT '
  'kiểu như giải phóng. Em đã giấu nó một thời gian dài. Em không cần sợ '
  'bố mẹ em sẽ bị tổn thương như thế nào, nhưng em cảm thấy thoải mái hơn '
  'như em cảm thấy tự tin hơn. Mọi người từng trêu em với những cái tên '
  'như ‘gà trống nam tính’ và những cái đó khiến em khó chịu. Nhưng khi em '
  'chia sẻ về em là ai, mọi người tôn trọng điều đó. Hơn nữa, khi em công '
  'khai, em đã tìm được nhiều bạn trong cộng đồng giống em, và em thấy có '
  'rất nhiều người tài năng trong cộng đồng LGBT." Các học sinh LGBTQ chia '
  'sẻ nhiều cảm xúc khác nhau về tương lai. Một số báo cáo cảm giác tích '
  'cực và phấn khởi về tương lai. Như một học sinh chia sẻ: "nhìn chung, '
  'em cũng khá tích cực về tương lai vì em đi học ở trường mà em thích, '
  'với bạn bè yêu thương em, để học một chuyên ngành mà em thích và cảm '
  'thấy tự tin về nó." Tuy nhiên, một học sinh khác bày tỏ sự bi quan về '
  'tương lai, nêu: "em không biết phải làm gì, tương lai khá mờ mịt."')

doc.add_page_break()

# ============================================================
# CHUONG 6 — CAC YEU TO LIEN QUAN DEN TRUONG HOC
# ============================================================
H('CHƯƠNG 6: CÁC YẾU TỐ LIÊN QUAN ĐẾN TRƯỜNG HỌC ẢNH HƯỞNG ĐẾN SỨC KHOẺ TÂM THẦN '
  'VÀ HẠNH PHÚC HỌC SINH TẠI VIỆT NAM', level=1)
P('(Chapter 6 — School-Related Factors Impacting Student Mental Health and Well-being in Viet Nam)',
  italic=True, align='center', size=11)

# --- Trang 44 ---
PageMark('--- Trang 44, UNICEF Việt Nam, 2022 ---')

H('Các phát hiện chính (Key Findings)', level=2)

P('1. Một số lượng đáng kể học sinh vị thành niên báo cáo mức độ gắn kết trường '
  'học thấp và ít kết nối với giáo viên. Học sinh thường không cảm thấy thoải '
  'mái khi đến với giáo viên để được hỗ trợ học tập hoặc cảm xúc – xã hội.', bold=True)

P('2. Các hiệu trưởng, nhà quản lý và chuyên gia cấp Bộ bày tỏ lo ngại về mối '
  'quan hệ giáo viên – học sinh. Các rào cản then chốt để cải thiện mối quan hệ '
  'giáo viên – học sinh bao gồm giáo viên thiếu thời gian và hỗ trợ hành chính.',
  bold=True)

P('3. Dữ liệu khảo sát học sinh chứng minh rằng học sinh cảm thấy ít gắn kết với '
  'trường học (ví dụ: ít kết nối với giáo viên) có tỷ lệ các vấn đề sức khoẻ tâm '
  'thần cao hơn. Điều quan trọng là các em gái báo cáo cảm thấy ít gắn kết với '
  'trường học hơn so với các em trai.', bold=True)

P('4. Bắt nạt là yếu tố nguy cơ nghiêm trọng đối với các vấn đề sức khoẻ tâm '
  'thần và hạnh phúc học sinh. Giáo viên không phải lúc nào cũng nhận ra tác '
  'động tiêu cực của bắt nạt lên học sinh.', bold=True)

P('5. Hầu hết tất cả người tham gia bày tỏ lo ngại nghiêm trọng về trải nghiệm '
  'áp lực học tập của học sinh và tác động của áp lực đó lên sức khoẻ tâm thần, '
  'hạnh phúc và việc học. Khối lượng công việc, áp lực từ giáo viên và phụ '
  'huynh, và các yếu tố hệ thống nằm sau mức áp lực học tập cao của học sinh.',
  bold=True)

P('6. Thiếu kiến thức sức khoẻ tâm thần của phụ huynh là yếu tố nguy cơ đối với '
  'các vấn đề sức khoẻ tâm thần học sinh.', bold=True)

P('7. Đại dịch COVID-19 đã có tác động tiêu cực lên sức khoẻ tâm thần và hạnh '
  'phúc học sinh. Tất cả người tham gia bày tỏ lo ngại về tác động của các hạn '
  'chế xã hội và cô lập lên sự phát triển tâm lý – xã hội học sinh, và tác '
  'động của giáo dục trực tuyến lên sức khoẻ tâm thần và việc học của học sinh.',
  bold=True)

P('8. Học sinh dân tộc thiểu số có nguy cơ tăng cao về các vấn đề sức khoẻ tâm '
  'thần và hạnh phúc. Đói nghèo gia đình, với áp lực liên quan lên học sinh '
  'phải làm việc (trong giờ học), kết hôn sớm và ở lại cộng đồng, là yếu tố '
  'nguy cơ cụ thể đối với học sinh dân tộc thiểu số. Thiếu kết nối gia đình – '
  'trường học và hiểu biết sức khoẻ tâm thần kém trong cộng đồng dân tộc thiểu '
  'số là các yếu tố nguy cơ bổ sung.', bold=True)

P('9. Học sinh LGBTQ có nguy cơ tăng cao về các vấn đề sức khoẻ tâm thần và '
  'đối mặt với các thách thức cụ thể, bao gồm kỳ thị, phân biệt đối xử, bắt '
  'nạt và các yếu tố stress trong mối quan hệ gia đình.', bold=True)

# --- Trang 45 ---
PageMark('--- Trang 45, UNICEF Việt Nam, 2022 ---')

P('Nhận thức của học sinh và giáo viên về khí hậu trường học, bao gồm môi '
  'trường trường học, gắn kết học sinh và an toàn trường học, đã gợi ý một số '
  'yếu tố nguy cơ quan trọng đối với các vấn đề sức khoẻ tâm thần vị thành '
  'niên. Xem Phụ lục 3 để biết các bảng và kết quả dữ liệu định lượng.')

H('6.1. Môi trường trường học (School Environment)', level=2)

P('Người trẻ nhìn chung báo cáo hài lòng với trường của mình. Môi trường '
  'trường học quan trọng đối với học sinh, với nhiều học sinh tự nêu ra các '
  'đặc điểm tích cực hoặc tiêu cực về môi trường. Nhiều học sinh bình luận về '
  '"cơ sở vật chất đẹp" của trường. Trong khi một học sinh báo cáo thích '
  'trường vì "môi trường xanh và ngăn nắp," các học sinh khác bày tỏ không '
  'hài lòng về mức độ ồn ào và sự hỗn loạn chung. Một học sinh nêu: "Em không '
  'thích các bạn học sinh gây ồn ào trong lớp và đến các lớp khác ngăn cản '
  'người khác học tập, mặc dù có quy tắc trường học để các bạn tuân theo '
  'nhưng các bạn vẫn như vậy." Tiếng ồn, sự hỗn loạn và việc thực thi quy tắc '
  'không nhất quán có tác động tiêu cực đến sự hài lòng của học sinh ở trường.')

P('Dữ liệu khảo sát học sinh định lượng chỉ ra rằng mức độ hài lòng thấp hơn '
  'với môi trường trường học có liên quan đáng kể đến tỷ lệ các vấn đề sức '
  'khoẻ tâm thần học sinh cao hơn. Không có khác biệt giữa các em trai và em '
  'gái nổi lên về nhận thức môi trường trường học. Học sinh ở Hà Nội hài lòng '
  'nhất với môi trường trường học so với học sinh từ các tỉnh khác.')

H('6.2. Gắn kết trường học và hỗ trợ giáo viên '
  '(School Engagement and Teacher Support)', level=2)

P('Gắn kết trường học bao gồm các yếu tố như kết nối của học sinh với giáo '
  'viên, kết nối học sinh, gắn kết học tập, kết nối toàn trường, văn hoá công '
  'bằng và tham gia của phụ huynh.')

P('Trong khi một số học sinh chia sẻ rằng các em cảm thấy được giáo viên hỗ '
  'trợ, phần lớn học sinh báo cáo cảm thấy không kết nối với giáo viên. Một '
  'số lo ngại tập trung vào sự thiếu hỗ trợ học tập. Ở Hà Nội, học sinh cảm '
  'thấy giáo viên thường chỉ rà soát nhanh tài liệu cơ bản trước khi tập '
  'trung vào nội dung nâng cao. Học sinh trong nhóm thảo luận về việc các em '
  'thường không hiểu những điều cơ bản và sau đó phải vật lộn tự rà soát nội '
  'dung đó bằng cách sử dụng internet. Khi được hỏi liệu các em có thể chia '
  'sẻ với giáo viên về việc cần nhiều hỗ trợ hơn, các em bày tỏ lo lắng về '
  'cách các em sẽ được giáo viên nhìn nhận.')

P('Khi được hỏi về việc đến với giáo viên để được hỗ trợ các vấn đề cảm xúc '
  '– xã hội, học sinh bày tỏ một số dè dặt. Các em không phải lúc nào cũng '
  'cảm thấy kết nối với giáo viên. Như một học sinh chia sẻ: "Giữa em và giáo '
  'viên, là một khoảng cách rất xa." Cũng có lo ngại rằng giáo viên sẽ không '
  'hiểu các vấn đề của các em và chỉ nói "Ừ, hãy mạnh mẽ lên." Một học sinh '
  'khác chia sẻ lo ngại rằng các em sẽ bị giáo viên nhìn nhận là yếu đuối, '
  'nêu: "Em nghĩ giáo viên sẽ nghĩ vấn đề này là buồn cười vì họ không hiểu '
  'tại sao những điều bình thường lại làm em cảm thấy áp lực." Và học sinh '
  'chia sẻ lo ngại về thời gian của giáo viên: "Giáo viên khá bận, nên chúng '
  'em không muốn làm phiền thầy cô với những việc riêng tư nhỏ nhặt của mình."')

P('Các hiệu trưởng, nhà quản lý cấp huyện và chuyên gia cấp Bộ thường chia sẻ '
  'lo ngại về gắn kết trường học của học sinh và các mối quan hệ với giáo '
  'viên. Một nhà quản lý Sở Y tế từ Hà Nội mô tả hai trường có khí hậu trường '
  'học rất khác nhau: "Ở một trường THPT, giáo viên cởi mở hơn. Học sinh có '
  'thể nói chuyện với giáo viên như bạn bè. Khi các em nói chuyện với giáo '
  'viên, các em cởi mở và tự do nói lên ý kiến của mình. Mặt khác, khi chúng '
  'tôi làm việc với [trường kia], học sinh đã phấn khích và nói chuyện cởi '
  'mở cho đến khi giáo viên vào. Khi giáo viên vào, các em im lặng. Các em '
  'không nói và giáo viên phải đặt câu hỏi và gọi tên để trả lời các câu hỏi '
  'đó. Các em quá ngại để bày tỏ ý kiến. Vì vậy... tôi nghĩ giáo viên và phụ '
  'huynh có tác động đến sức khoẻ tâm thần học sinh."')

# --- Trang 46 ---
PageMark('--- Trang 46, UNICEF Việt Nam, 2022 ---')
P('Một số nhà quản lý chia sẻ lo ngại rằng giáo viên thiếu thời gian và hỗ '
  'trợ để phát triển các mối quan hệ với học sinh. Như một nhà quản lý từ Sở '
  'LĐ-TB&XH Hà Nội nêu: "Khi tôi nghĩ lại thời của chúng tôi, giáo viên có '
  'thể biết tình hình gia đình của mỗi học sinh rất rõ. Nhưng bây giờ, nếu '
  'tôi nói về các giáo viên chủ nhiệm, tôi nghĩ đó sẽ là yêu cầu quá lớn cho '
  'giáo viên vì áp lực dạy học quá lớn." Các nhà quản lý khác chia sẻ rằng '
  'vấn đề này bị khuếch đại bởi sĩ số lớp lớn khiến giáo viên khó làm quen '
  'với từng học sinh.')

P('Dữ liệu khảo sát học sinh cung cấp bằng chứng rằng học sinh cảm thấy ít '
  'gắn kết với trường học (tức là ít kết nối với giáo viên, với các học sinh '
  'khác, ít kết nối toàn trường) có tỷ lệ các vấn đề sức khoẻ tâm thần cao '
  'hơn. Điều quan trọng là các em gái báo cáo cảm thấy ít gắn kết với trường '
  'học hơn so với các em trai. Học sinh trung học cơ sở báo cáo mức gắn kết '
  'trường học cao hơn so với học sinh trung học phổ thông. Học sinh từ Điện '
  'Biên báo cáo ít gắn kết trường học nhất, trong khi học sinh từ Đồng Tháp '
  'gắn kết nhiều nhất.')

H('6.3. Các mối quan hệ bạn bè và bắt nạt (Peer Relationships and Bullying)',
  level=2)

P('Học sinh nhìn chung báo cáo cảm thấy an toàn và có tình bạn tốt ở trường. '
  'Các em thừa nhận rằng các tranh cãi và vấn đề có thể xảy ra giữa bạn bè, '
  'nhưng chia sẻ các chiến lược để quản lý những vấn đề đó, bao gồm xin lỗi '
  'và dành thời gian để nguôi giận trước khi "làm hoà." Học sinh nói về tầm '
  'quan trọng của tình bạn và chia sẻ các khó khăn trong việc điều hướng các '
  'vấn đề với tình bạn. Một học sinh nói: "Chúng em đang trong giai đoạn phát '
  'triển tâm lý liên tục, nên ngay cả một vấn đề nhỏ hoặc một hiểu lầm giữa '
  'bạn bè cũng có thể gây ra sự đổ vỡ trong tình bạn. Vì vậy, mỗi lần tình '
  'bạn đổ vỡ là một điều rất khó vượt qua. Đôi khi nước mắt có thể xuất '
  'hiện, đặc biệt khi chúng em tin vào tình bạn." Khi xung đột bạn bè xảy '
  'ra, chúng có tác động tiêu cực lên trải nghiệm của học sinh ở trường. '
  'Như một học sinh ở Hà Nội chia sẻ: "Có một số bạn liên tục có xung đột '
  'với nhau. Nó khiến lớp mất đi sự đoàn kết và khiến chúng em những người '
  'không tham gia vào đó, cảm thấy căng thẳng hơn."')

P('Giáo viên đồng ý với học sinh rằng các mối quan hệ bạn bè là yếu tố quan '
  'trọng cho hạnh phúc học sinh. Giống như học sinh, giáo viên chia sẻ rằng '
  'một số học sinh có xung đột bạn bè dẫn các em rút lui khỏi các mối quan '
  'hệ khác và khỏi trường học, và cảm thấy bị cô lập và chán nản.')

P('Học sinh và giáo viên đồng ý rằng các cuộc đánh nhau thể chất giữa học '
  'sinh hiếm khi xảy ra ở trường. Trong khi học sinh có vẻ nhìn chung cảm '
  'thấy an toàn về thể chất ở trường, học sinh có thể không phải lúc nào '
  'cũng cảm thấy an toàn về mặt cảm xúc. Học sinh chia sẻ nhiều ví dụ về '
  'bắt nạt giữa bạn bè. Một học sinh nêu rằng một bạn cùng lớp đã "bị tẩy '
  'chay và bị mọi người trong lớp bỏ rơi" trong nhiều tháng. Sau một thời '
  'gian, một giáo viên khuyên các học sinh nên tử tế và bao dung với học '
  'sinh bị nhắm đến. Sự hỗ trợ không chính thức này từ giáo viên được học '
  'sinh tham gia đánh giá cao. Nhưng cô ấy báo cáo rằng (a) nhiều học sinh '
  'không cảm thấy thoải mái khi chia sẻ với giáo viên về các khó khăn xã '
  'hội hoặc cảm xúc và (b) không có hỗ trợ nào khác của trường được cung '
  'cấp để giải quyết bắt nạt hoặc hỗ trợ nạn nhân. Học sinh từ nhiều trường '
  'cũng báo cáo loại "bắt nạt tẩy chay" này.')

P('Kết quả khảo sát học sinh chỉ ra rằng nhận thức của học sinh về trường '
  'học là ít an toàn hơn có liên quan đáng kể đến tỷ lệ các vấn đề sức khoẻ '
  'tâm thần cao hơn. Dữ liệu cũng cho thấy tỷ lệ bắt nạt cao hơn ở trung học '
  'cơ sở so với trung học phổ thông. Học sinh ở Hà Nội báo cáo nhận thức '
  'trường học an toàn cao hơn so với học sinh ở ba tỉnh khác. Học sinh ở '
  'Gia Lai báo cáo an toàn trường học ít đáng kể hơn so với học sinh ở các '
  'tỉnh khác.')

# --- Trang 47 ---
PageMark('--- Trang 47, UNICEF Việt Nam, 2022 ---')
P('Giáo viên không đồng ý với nhau về tần suất và tác động của bắt nạt. Một '
  'số giáo viên báo cáo rằng có rất ít bắt nạt ở trường, trong khi những '
  'người khác thấy nó xảy ra thường xuyên. Một giáo viên ở Hà Nội chia sẻ '
  'câu chuyện về xung đột bạn bè trong một lớp liên quan đến bắt nạt. Cô '
  'tin rằng giáo viên và phụ huynh đã can thiệp thành công để giải quyết '
  'vấn đề. Cô báo cáo rằng những trải nghiệm như vậy là bình thường ở tuổi '
  'thơ, và trẻ em thường học hỏi và phát triển từ những trải nghiệm này.')

P('Học sinh đặc biệt khó chịu về bắt nạt tập trung vào ngoại hình của một '
  'học sinh. Các em nói rằng học sinh thường bị chế giễu vì "mập" hoặc '
  '"thấp" với các em gái thường bị nhắm đến nhiều nhất. Một học sinh mô tả '
  'trải nghiệm: "Kiểu như trêu chọc bạn cho đến khi bạn tức giận. Đến mức '
  'bạn khóc. Rồi cảm thấy bất lực. Bạn cảm thấy như bạn không thể làm gì '
  'được về chuyện đó." Các em báo cáo rằng trong khi cả các em trai và em '
  'gái đều trêu chọc các em gái theo cách này, chủ yếu là các em trai làm '
  'loại trêu chọc này. Khi được hỏi tại sao các em nghĩ các em trai trêu '
  'chọc các em gái theo cách này, các em nghĩ các em trai thích làm tổn '
  'thương cảm xúc của các em gái, nêu: "Kiểu như, họ cười lớn khi thấy ai '
  'đó buồn. Họ cười trên nỗi buồn của người khác."')

P('Bắt nạt mạng (cyberbullying) cũng là một vấn đề nghiêm trọng đối với học '
  'sinh. Tất cả học sinh báo cáo sử dụng internet và mạng xã hội, và nhiều '
  'học sinh báo cáo quan sát cả bạn bè lẫn người lạ bị bắt nạt trực tuyến. '
  'Dữ liệu khảo sát học sinh đã khám phá vấn đề bắt nạt mạng với mục: "Bạn '
  'đã từng bị bắt nạt mạng chưa?" Hầu hết học sinh (47,3 %) báo cáo rằng các '
  'em chưa bao giờ bị bắt nạt mạng. Khoảng 30 % học sinh chỉ "hiếm khi" bị '
  'bắt nạt mạng, 20,2 % báo cáo "đôi khi" và 2,1 % "thường xuyên" bị bắt '
  'nạt mạng. Không có khác biệt giới tính hoặc khác biệt tỉnh đáng kể nào '
  'được phát hiện. Tần suất bắt nạt mạng và đau khổ liên quan khác biệt '
  'đáng kể theo cấp học. Học sinh trung học phổ thông có khả năng trải qua '
  'bắt nạt mạng đáng kể hơn và báo cáo nhiều đau khổ liên quan đến bắt nạt '
  'mạng đáng kể hơn so với học sinh trung học cơ sở.')

P('Các hiệu trưởng và nhà quản lý giáo dục chia sẻ lo ngại về bắt nạt mạng. '
  'Một hiệu trưởng ở Hà Nội mô tả rằng học sinh "gọi tên nhau, đánh nhau '
  'hoặc xúc phạm nhau hoặc kéo bè phái trực tuyến chống lại nhau. Một ‘bè '
  'phái trực tuyến’ là khi học sinh của chúng ta tụ tập trực tuyến cùng '
  'nhau để bôi nhọ, xúc phạm, chỉ nhắm vào một người. Và tất cả những điều '
  'đó ảnh hưởng đến tâm lý của cả người bị tấn công lẫn những người tích '
  'cực tấn công."')

H('6.4. Áp lực trường học (School pressure)', level=2)

H('6.4.1. Áp lực học tập (Academic Pressure)', level=3)

P('Tất cả học sinh báo cáo trải nghiệm áp lực học tập và các FGD đã khám '
  'phá bản chất và tác động của áp lực học tập. Một vài học sinh tin rằng '
  'áp lực có ích trong việc tạo động lực cho các em học tập với một học '
  'sinh nói: "Em nghĩ nếu không có stress đi kèm với học tập, em đã không '
  'có sự chú ý cần thiết để tập trung vào học đạt điểm cao nhất." Tuy '
  'nhiên, hầu hết học sinh mô tả bị choáng ngợp bởi áp lực và stress và '
  'tin rằng áp lực có tác động tiêu cực lên sức khoẻ và thành tích học '
  'tập của các em. Học sinh bày tỏ mức stress cao liên quan đến khối '
  'lượng công việc, áp lực từ giáo viên và phụ huynh, thi cử và các lớp '
  'học thêm.')

P('Hầu hết các giáo viên tham gia FGD bày tỏ lo ngại về lượng áp lực học '
  'tập mà học sinh trải qua, và tác động của áp lực đó lên sức khoẻ tâm '
  'thần và hạnh phúc của các em. Giáo viên báo cáo các nguồn gây áp lực '
  'học tập bao gồm khối lượng công việc, áp lực từ các cuộc thi và bài '
  'thi, giáo viên, phụ huynh và tự kỳ vọng của học sinh. Giáo viên lưu ý '
  'sự khác biệt về tuổi trong nguyên nhân gây stress dựa trên trường học. '
  'Họ báo cáo rằng học sinh tuổi trẻ hơn có lo âu liên quan đến việc '
  'thích nghi với môi trường trường học mới, bao gồm nhiều lớp học và '
  'công việc, và việc cố gắng kết bạn mới. Học sinh lớn tuổi hơn trải qua '
  'áp lực liên quan đến việc đưa ra các quyết định cuộc sống quan trọng '
  'và chuẩn bị cho các bài thi chuyên.')

P('Giới tính là một yếu tố đáng kể trong áp lực học tập được học sinh báo '
  'cáo. Kết quả khảo sát chỉ ra rằng các em gái trải qua mức áp lực học '
  'tập cao hơn đáng kể bao gồm áp lực học tập, lo lắng về điểm, tự kỳ vọng '
  'và chán nản về học tập.')

# --- Trang 48 ---
PageMark('--- Trang 48, UNICEF Việt Nam, 2022 ---')
P('Điều này bất chấp nhận thức tương tự về khối lượng công việc học tập '
  'giữa các em trai và em gái. Cấp học cũng tác động đến các trải nghiệm '
  'áp lực học tập. Các phát hiện gợi ý rằng học sinh trung học phổ thông '
  'báo cáo áp lực học tập đáng kể hơn so với học sinh trung học cơ sở. '
  'Stress học tập đến từ áp lực học tập, tự kỳ vọng, khối lượng công '
  'việc và trải nghiệm chán nản lớn hơn của các em. Học sinh từ các tỉnh '
  'khác nhau cho thấy sự khác biệt đáng kể trong trải nghiệm áp lực học '
  'tập. Vị thành niên báo cáo mức áp lực học tập cao hơn ở Gia Lai và Hà '
  'Nội so với Đồng Tháp và Điện Biên.')

H('6.4.2. Khối lượng công việc và áp lực học tập '
  '(Workload and Academic Pressure)', level=3)

P('Một nguyên nhân của stress học tập là khối lượng công việc cần thiết '
  'để đạt điểm tốt. Hầu hết tất cả học sinh, từ cả khu vực nông thôn và '
  'đô thị, đồng ý rằng khối lượng công việc là quá nặng. Như một học sinh '
  'từ Hà Nội chia sẻ đùa: "Em có một quy luật gọi là ‘Quy luật Bài tập '
  'Vô tận.’" Em giải thích: "Em không thể có nhiều giải trí. Em có các '
  'lớp học thêm, rồi mỗi sáng Thứ Bảy em phải làm bài tập về nhà của cả '
  'tuần. Chủ Nhật em cũng có các lớp học thêm."')

P('Khối lượng công việc cao được phản ánh trong các phản hồi khảo sát '
  'của học sinh về thời gian dành cho học tập ngoài trường học. Kết quả '
  'chỉ ra rằng hơn 50 % học sinh học hơn 2 giờ sau giờ học mỗi ngày, với '
  '28 % học sinh báo cáo hơn 3 giờ học mỗi đêm. Học sinh cũng dành một '
  'lượng thời gian đáng kể cho việc học thêm hoặc các lớp phụ mỗi tuần. '
  'Trong khi khoảng 50 % học sinh báo cáo 2 giờ hoặc ít hơn các lớp học '
  'thêm hoặc học thêm mỗi tuần, 50 % học sinh còn lại báo cáo hơn 3 giờ '
  'mỗi tuần, bao gồm 15 % học sinh dành hơn 9 giờ mỗi tuần cho các lớp '
  'học thêm/dạy kèm.')

P('Giáo viên chia sẻ lo ngại của học sinh về khối lượng công việc nặng '
  'và sự thiếu thời gian cho các hoạt động khác. Nhiều giáo viên cảm '
  'thấy rằng khối lượng công việc học sinh cao ngăn các em học các kỹ '
  'năng quan trọng. Một giáo viên ở Đồng Tháp chia sẻ: "Ngày nay kỹ '
  'năng sống và kỹ năng mềm của học sinh không tốt như của chúng tôi. '
  'Các em không thể bơi ở sông giỏi như chúng tôi. Tôi không biết phép '
  'so sánh này có hợp lý không, nhưng thực tế là học sinh ngày nay '
  'giống như gà công nghiệp. Chúng tôi bắt các em học tất cả thời gian '
  'và đó là tất cả các em biết. Các em không có ý tưởng nào về những '
  'thứ khác."')

H('6.4.3. Áp lực học tập từ giáo viên (Academic Pressure from Teachers)',
  level=3)

P('Nhiều học sinh bày tỏ lo âu về phản ứng của giáo viên đối với thành '
  'tích kém của học sinh. Khi học sinh không thể hiện tốt, các em báo '
  'cáo rằng giáo viên "không vui," có thể "la hét," và "đưa ra so '
  'sánh." Học sinh dường như đặc biệt đau khổ khi giáo viên so sánh các '
  'em với những học sinh khác thể hiện tốt hơn, nêu rằng điều này cảm '
  'thấy "không công bằng" và khiến các em cảm thấy "buồn." Một học sinh '
  'giải thích phản ứng của mình đối với điều này bằng cách nói: "giáo '
  'viên thường đặt các câu hỏi khó thuộc về các lớp năng khiếu, và khi '
  'bạn không thể trả lời điều đó, nó dẫn đến so sánh với các lớp khác '
  'và dẫn đến học sinh cảm thấy căm ghét giáo viên đó, và khi lớp của '
  'họ đến, không học sinh nào muốn giơ tay." Phản ứng tiêu cực của '
  'giáo viên đối với thành tích của học sinh có thể làm giảm gắn kết '
  'và thành công học tập của học sinh. Như một học sinh nêu: "Khi '
  'chúng em bị stress, chúng em không cảm thấy tốt, hoặc cảm thấy lo '
  'âu khi bạn phải đọc thuộc bài cho giáo viên vì bạn sợ có câu trả '
  'lời sai và làm giáo viên không vui."')

H('6.4.4. Áp lực học tập từ phụ huynh (Academic Pressure from Parents)',
  level=3)

P('Học sinh cũng thảo luận về stress áp lực học tập từ phụ huynh. Nhiều '
  'học sinh báo cáo rằng phụ huynh gây áp lực để các em thành công và '
  'có tiêu chuẩn cao được báo cáo là gây stress cho học sinh phải "học '
  '24/24." Như một học sinh mô tả: "Bố mẹ em không ép em, nhưng nhiều '
  'lúc bố mẹ em có kỳ vọng. Vì vậy thường khi em có điểm thấp bố mẹ em '
  'không mắng em nhưng em cảm thấy mình làm họ thất vọng. Những điều đó '
  'khiến em bị stress."')

doc.add_page_break()
# --- Trang 49 — Case study Two Schools ---
PageMark('--- Trang 49, UNICEF Việt Nam, 2022 ---')
H('HỘP CASE STUDY — HAI TRƯỜNG, HAI CÁCH TIẾP CẬN HỌC TẬP: '
  'So sánh các mô hình trường học và hạnh phúc học sinh', level=2)
P('(Two Schools, Two Approaches to Learning: A Comparison of School Models '
  'and Student Well-Being)', italic=True, size=11)

P('Giáo viên từ hai trường ở Hà Nội đã chia sẻ các quan điểm chỉ ra các nhận '
  'thức và cách tiếp cận khác nhau đối với áp lực học tập và khối lượng công '
  'việc.')

P('Trường 1 — Quan tâm đến sức khoẻ tâm thần:', bold=True)
P('Giáo viên ở một trường bày tỏ lo ngại về áp lực học tập, chia sẻ: "Giáo '
  'viên thực sự quan tâm đến sức khoẻ tâm thần của học sinh vì họ hiểu rằng '
  'học sinh không thể học tốt khi cảm thấy bị stress và tôi đặc biệt không '
  'muốn các em cảm thấy bị áp lực." Giáo viên báo cáo giao ít bài tập về nhà '
  'cho phép học sinh rà soát và khám phá các chủ đề lớp học nhưng không làm '
  'các em choáng ngợp với công việc. Họ mô tả suy nghĩ sáng tạo về bài tập '
  'về nhà. Ví dụ, một giáo viên Hoá học nói về "các bài tập liên quan đến '
  'các thí nghiệm khuyến khích các em theo dõi bằng chứng và quan sát các '
  'hiện tượng đời thực" mà học sinh có thể chia sẻ với lớp qua video. Tuy '
  'nhiên, giáo viên ở trường này bày tỏ lo ngại về lượng học tập mà học sinh '
  'làm, và tự hỏi liệu nó có thể cuối cùng có tác động tiêu cực đến việc học '
  'không. Như một giáo viên nêu: "Tôi nhớ đã học ít hơn nhiều so với học sinh '
  'ngày nay và tôi thực sự chơi hầu hết thời gian. Tôi cảm thấy như cân bằng '
  'thời gian của mình giữa học tập và giải trí chắc chắn khiến tôi hấp thụ '
  'kiến thức tốt hơn và hiệu quả hơn. Tuy nhiên, khi tôi nhìn học sinh ngày '
  'nay, nhiều em được thấy học suốt thời gian, từ sáng sớm đến tối khuya." '
  'Giáo viên ở trường này tin rằng cấu trúc, chương trình giảng dạy và cách '
  'tiếp cận giảng dạy tại trường của họ giảm stress học sinh. Học sinh ở các '
  'lớp nhỏ không được tổ chức theo thành tích hay xếp hạng nên có ít cạnh '
  'tranh giữa học sinh hơn. Trường nuôi dưỡng điểm mạnh của các em và khuyến '
  'khích các em học và chuyên sâu vào các môn các em giỏi. Trường cung cấp '
  'các câu lạc bộ và hoạt động để giúp học sinh trau dồi sở thích và kỹ năng, '
  'bao gồm lãnh đạo, làm việc nhóm và tổ chức.')

P('Trường 2 — Chấp nhận áp lực học tập:', bold=True)
P('Cách tiếp cận đối với áp lực học tập và khối lượng công việc này khác với '
  'cách tiếp cận của một trường Hà Nội khác. Giáo viên trong nhóm tập trung '
  'từ trường này không xem áp lực học tập hoặc stress là một vấn đề. Giáo '
  'viên ở trường này tin rằng chỉ một số ít học sinh bị ảnh hưởng đáng kể '
  'bởi stress. Như một giáo viên nêu: "Khi tình huống đó xảy ra, giáo viên, '
  'gia đình và bạn bè sẽ giúp các em nhanh chóng để tránh tình huống đó và '
  'giúp các em nhanh chóng trở về trạng thái vui vẻ, bình thường." Giáo viên '
  'ở đây không thấy stress ảnh hưởng đến việc học nhiều. Như một giáo viên '
  'mô tả: "Về khả năng học tập, [stress] sẽ có ít tác động, tác động rất '
  'nhỏ, và hầu như không đáng nói, và học sinh sẽ vượt qua nó nhanh chóng '
  'và không bị ảnh hưởng nhiều." Giáo viên nhận thức được khối lượng công '
  'việc lớn đặt lên học sinh ở trường này và coi điều này là cần thiết cho '
  'học sinh ở một trường cạnh tranh. Như một giáo viên nêu: "Về trường của '
  'tôi, đó là một trong những trường top của thành phố, và tất nhiên, nếu '
  'các em muốn xuất sắc ở trường thì cường độ và khối lượng công việc chắc '
  'chắn sẽ phải nhiều hơn so với các trường khác. Thời gian dành cho học '
  'tập rất nhiều, gần như cả tuần, ngay cả Thứ Bảy và Chủ Nhật các em vẫn '
  'phải ở nhà. Các em hiếm khi đi ra ngoài và vào cuối tuần các em dùng '
  'thời gian để tham gia các lớp học thêm."')

P('Giáo viên ở đây thấy một số trẻ em vật lộn, nhưng không nhất thiết xác '
  'định rằng điều này có thể do stress hoặc các vấn đề sức khoẻ tâm thần '
  'gây ra, hoặc rằng stress và thất bại có thể dẫn đến sức khoẻ tâm thần '
  'tệ hơn. Ví dụ, một giáo viên báo cáo: "Từ những gì tôi thấy, ví dụ '
  'trong cuộc thi Học sinh Giỏi, các em rất stress, lo lắng, cảm thấy áp '
  'lực về bài thi này, lo lắng về tỷ lệ học sinh thắng kỳ thi này. Kiến '
  'thức các em phải học là rất lớn, nhưng các em cực kỳ quyết tâm. Những '
  'em có quyết tâm mạnh và nỗ lực lớn, các em biến stress thành động lực '
  'để vượt qua giai đoạn đó một cách xuất sắc và ngoạn mục. Nhưng những '
  'em khác có năng lực cá nhân yếu hơn, khả năng tự kiểm soát của các em '
  'kém. Đôi khi sau cả những thất bại nhỏ rồi các em đột ngột thất bại."')

P('So sánh giữa hai trường này chứng minh hai cách tiếp cận rất khác nhau '
  'đối với chương trình giảng dạy, mối quan hệ giáo viên – học sinh và '
  'các mục tiêu giáo dục. Rõ ràng là văn hoá trường học của trường đầu '
  'tiên đặt các mối quan hệ học sinh – giáo viên và hạnh phúc học sinh ở '
  'trung tâm chương trình giảng dạy trường học. Cách tiếp cận này có khả '
  'năng cung cấp hỗ trợ cảm xúc – xã hội lớn hơn cho nhiều học sinh hơn, '
  'qua đó cho phép nhiều học sinh gắn kết và thành công một cách tích '
  'cực ở trường và xa hơn nữa.', italic=True)

doc.add_page_break()

# --- Trang 50 ---
PageMark('--- Trang 50, UNICEF Việt Nam, 2022 ---')
P('(Tiếp 6.4.4 Áp lực học tập từ phụ huynh.)')
P('Một học sinh khác nêu: "Em thường bị stress do điểm số vì bố mẹ em muốn '
  'em đạt một điểm nhất định và lấy danh hiệu Xuất sắc." Một học sinh ở Hà '
  'Nội nói: "Em không có nhiều thời gian để nghỉ ngơi, vì em đang thực sự '
  'yếu môn toán, nên sau khi tan học bố em luôn bắt em đi làm toán." Cả '
  'bốn học sinh lớp 8 trong nhóm học sinh Điện Biên báo cáo cảm thấy áp '
  'lực từ gia đình về các kế hoạch tương lai của các em. Các học sinh chia '
  'sẻ rằng các em "mơ ước trở thành một cái gì đó, rồi cố gắng" nhưng cảm '
  'nhận rằng bố mẹ không quan tâm đến ước mơ của các em. Phụ huynh, từ quan '
  'điểm của các em, "dẫn chúng em làm những gì họ muốn" và trẻ em phải '
  '"học để trở thành những gì bố mẹ muốn." Các em chia sẻ rằng phụ huynh '
  'muốn các em "làm các công việc kiểu làm-việc-văn-phòng" ở "những nơi '
  'mà chúng em có thể dễ dàng vào vì bố mẹ có mối quan hệ ở đó." Không '
  'may, các học sinh đồng ý rằng những nơi này thường là "những nơi chúng '
  'em không thích."')

H('6.4.5. Nguyên nhân hệ thống của áp lực học tập '
  '(Systemic Causes of Academic Pressure)', level=3)

P('Các nhà quản lý và chuyên gia từ các ngành giáo dục, xã hội và y tế đã '
  'thảo luận các nguyên nhân hệ thống nền tảng của áp lực học tập. Hầu hết '
  'tất cả người cung cấp thông tin chủ chốt bày tỏ lo ngại về tác động '
  'của áp lực học tập lên sức khoẻ tâm thần và hạnh phúc vị thành niên. '
  'Một chuyên gia từ Sở LĐ-TB&XH Hà Nội lưu ý rằng áp lực học tập đến từ '
  'phụ huynh, giáo viên và học sinh, và lưu ý rằng chương trình giáo dục '
  'không dạy trẻ em cách giải quyết các vấn đề, bao gồm cách đối phó với '
  'stress trường học. Các vấn đề hệ thống nền tảng được mô tả là một '
  '"bệnh thành tích" trong hệ thống trường học. Một chuyên gia từ Bộ Y '
  'tế đồng ý và bày tỏ lo ngại về chương trình giảng dạy trường học. '
  'Chuyên gia này tin rằng chương trình giảng dạy quá tập trung vào '
  'chuẩn bị thi cường độ cao ở tuổi trẻ và các cuộc thi học sinh. Khi mô '
  'tả cách các cuộc thi tác động đến khối lượng công việc, áp lực và chất '
  'lượng giáo dục của học sinh, ông nêu: "Rõ ràng là một hệ thống dựa '
  'trên thành tích. Nó mang lại danh tiếng tốt cho các trường và đặt áp '
  'lực lên học sinh."')

P('Giáo viên và các nhà quản lý khám phá các nguyên nhân nền tảng của áp '
  'lực giáo viên. Các nhà quản lý Sở GD&ĐT lưu ý rằng hệ thống đánh giá '
  'thành tích giáo viên dẫn đến tăng áp lực học tập lên học sinh. Thành '
  'tích giáo viên dựa trên thành tích học sinh và hệ thống bao gồm cạnh '
  'tranh giữa các giáo viên cho các phần thưởng và thăng tiến giáo viên. '
  'Như một chuyên gia Sở GD&ĐT lưu ý: "Hãy tưởng tượng kịch bản khi giáo '
  'viên được yêu cầu có học sinh đạt thành tích cao để nâng hạng lớp học '
  'và rồi một số lớp bị xếp hạng thấp hơn chỉ vì một số học sinh không '
  'làm được và sau đó ai đó đánh giá bạn và thành tích của bạn. Điều đó '
  'thực sự khó khăn." Tỷ lệ ghi danh trường học cao hơn góp phần vào vấn '
  'đề này vì nhiều học sinh hơn có nghĩa là khó hơn để học sinh và giáo '
  'viên đạt thứ hạng cao nhất, đặt cả học sinh và giáo viên dưới áp lực '
  'tăng cao.')

H('6.4.6. Tác động của áp lực học tập: Sức khoẻ tâm thần và hạnh phúc '
  '(Impact of Academic Pressure: Mental Health and Well-Being)', level=3)

P('Học sinh cung cấp rất nhiều bằng chứng về tác động tiêu cực của áp lực '
  'học tập, bao gồm stress và các vấn đề sức khoẻ tâm thần khác, tập trung '
  'và trí nhớ kém, thiếu ngủ và các vấn đề xã hội. Mối quan ngại chính là '
  'mối quan hệ giữa áp lực học tập và sức khoẻ tâm thần học sinh. Các phân '
  'tích tương quan từ dữ liệu khảo sát học sinh chỉ ra mối quan hệ có ý '
  'nghĩa cao, trung bình – mạnh giữa các trải nghiệm áp lực học tập của '
  'học sinh và sức khoẻ tâm thần được báo cáo. Áp lực học tập, lo lắng về '
  'điểm, chán nản liên quan đến học tập, tự kỳ vọng và khối lượng công '
  'việc đều có liên quan đáng kể với tất cả các vấn đề sức khoẻ tâm thần. '
  'Có mối quan hệ mạnh giữa tổng áp lực học tập và tổng khó khăn sức khoẻ '
  'tâm thần. Vì các phân tích tương quan không thể chứng minh chiều hướng '
  'của mối quan hệ, có thể là (a) áp lực học tập có thể ảnh hưởng đến sức '
  'khoẻ tâm thần học sinh, (b) sức khoẻ tâm thần học sinh có thể ảnh '
  'hưởng đến các trải nghiệm áp lực học tập của học sinh, (c) cả hai có '
  'thể ảnh hưởng lẫn nhau hoặc (d) điều gì đó khác có thể ảnh hưởng đến '
  'cả hai.')

# --- Trang 51 ---
PageMark('--- Trang 51, UNICEF Việt Nam, 2022 ---')
P('Học sinh ở Hà Nội báo cáo mức áp lực học tập cao và chia sẻ nhiều ví '
  'dụ về tác động tiêu cực của áp lực học tập lên sức khoẻ tâm thần và '
  'hạnh phúc. Một học sinh chia sẻ: "Khi chúng em quá stress, chúng em '
  'không có thời gian để tập thể dục hoặc đi bộ hoặc thư giãn. Ngay cả '
  'khi không phải vì dịch, chúng em cũng không thể đi đâu thoải mái. Bởi '
  'vì bây giờ bất cứ nơi nào em đi em cũng phải nghĩ về bài tập. Em nghĩ '
  'về nó mọi lúc — trường học, bài tập, và mọi thứ. Khi có quá nhiều '
  'công việc và em không thể làm kịp hoặc thiếu gì đó, nó khá đáng sợ. '
  'Em sợ bị giáo viên hoặc bố mẹ mắng. Khi em bị stress, em không thể '
  'ngủ yên. Ví dụ, nếu đó là trước kỳ thi, khi em học môn văn, khi em '
  'ngủ đầu em vẫn đang nghĩ về văn. Em không thể ngủ."')

P('Giáo viên chia sẻ một số ví dụ về áp lực học tập dẫn đến sức khoẻ tâm '
  'thần kém và các vấn đề chức năng. Giáo viên ở Đồng Tháp nói về mối '
  'liên kết giữa áp lực học tập và sức khoẻ tâm thần, và tin rằng các kỳ '
  'vọng học tập nên được điều chỉnh theo khả năng của từng học sinh để '
  'gắn kết thành công các em với trường học. Một giáo viên ở Hà Nội chia '
  'sẻ: "Năm ngoái em có một học sinh chuyên về Hoá. Ban đầu, các em có '
  'hy vọng cao cho Cuộc thi Học sinh Giỏi nhưng do áp lực nặng nề các '
  'em đặt lên mình trong giai đoạn ôn tập, nó phản tác dụng, và các em '
  'quá stress. Các em bị stress nặng và không thể vượt qua. Phải mất '
  'phụ huynh và giáo viên rất nhiều thời gian và nỗ lực trước kỳ thi '
  'tuyển sinh vào trung học phổ thông để giúp các em tìm ra các em muốn '
  'đi đâu."')

P('Có sự lo ngại rộng rãi giữa các nhà quản lý và chuyên gia về tác động '
  'của áp lực học tập lên sức khoẻ tâm thần và hạnh phúc học sinh. Một '
  'nhà quản lý từ Sở LĐ-TB&XH Hà Nội bày tỏ lo ngại về gánh nặng của áp '
  'lực học tập và khối lượng công việc nêu: "Trẻ em chịu quá nhiều áp '
  'lực nên các em lo lắng, các em bị stress, các em không thể ăn, và '
  'các em không thể ngủ. Đó là lý do tại sao rất nhiều học sinh — tôi '
  'biết có rất nhiều em — rất nhiều học sinh sau kỳ thi rất trầm cảm. '
  'Thực sự không đáng cho các em, tôi nghĩ." Một chuyên gia từ Bộ GD&ĐT '
  'đồng ý với mối quan ngại này và chia sẻ trải nghiệm cá nhân về con '
  'của một người bạn là học sinh giỏi nhưng khi đến lúc thi tuyển sinh '
  'đại học thì cô gặp khó khăn khi học.')

P('Các chuyên gia lo ngại về mức áp lực học tập cao đặt lên học sinh ở '
  'các trường cạnh tranh. Một chuyên gia từ Sở GD&ĐT Hà Nội chia sẻ '
  'rằng học sinh ở một trường cạnh tranh địa phương bị "đưa vào danh '
  'sách đen" nếu các em không hoàn thành tất cả bài tập về nhà, nghĩa '
  'là những học sinh đó bị báo cáo cho trường và phụ huynh. Ông hỏi: '
  '"Bạn có nghĩ đó không phải là áp lực quá nhiều không? Sẽ ổn nếu học '
  'sinh có thể quản lý tất cả điều đó nhưng trong đời thực ngay cả học '
  'sinh xuất sắc cũng có những khoảnh khắc không hài lòng và thất bại."')

P('Học sinh cũng chia sẻ cách áp lực học tập và stress liên quan tác '
  'động tiêu cực đến việc học và thành tích học tập. Học sinh thảo luận '
  'cách các em đôi khi học rất chăm chỉ chỉ để phát hiện các em không '
  'có ký ức về thông tin khi đến lúc thi. Học sinh trung học phổ thông '
  'báo cáo mức stress học tập cao hơn. Nhiều học sinh báo cáo thức đến '
  '2 hoặc 3 giờ sáng những đêm trước kỳ thi.')

# --- Trang 52 ---
PageMark('--- Trang 52, UNICEF Việt Nam, 2022 ---')
P('Nhiều học sinh báo cáo rằng các em không ngủ đủ. Một học sinh chia sẻ: '
  '"Bình thường, em ngủ từ 2 giờ sáng đến 4 giờ sáng, hoặc 1 giờ sáng đến '
  '5 giờ sáng. Vì vậy, trong lớp em không thể tập trung. Em quá mệt để '
  'viết. Tình hình ngày càng tệ hơn. Em thực sự cảm thấy như em không '
  'thể theo kịp mọi thứ." Em nêu em đang nhận thấy khó khăn trong việc '
  'tập trung trong lớp và ghi nhớ mọi thứ. Một học sinh trung học cơ sở '
  'từ Đồng Tháp lưu ý: "Khi em bị stress, em thường mệt mỏi và em không '
  'muốn nói nhiều với người khác, điều này khiến các bạn khác trong lớp '
  'trêu em và em không thích điều đó."')

H('6.4.7. Tác động của áp lực học tập: Việc học của học sinh '
  '(Impact of Academic Pressure: Student Learning)', level=3)

P('Cùng với lo ngại về sức khoẻ tâm thần học sinh, người tham gia bày tỏ '
  'lo ngại rằng áp lực học tập có tác động tiêu cực lên việc học của học '
  'sinh. Một nhà quản lý Sở LĐ-TB&XH từ Gia Lai thấy rằng các trường bây '
  'giờ chỉ quan tâm đến thành tích học tập, có tác động tiêu cực đến sở '
  'thích nội tại của học sinh trong các môn học và đam mê học tập. Bà nêu: '
  '"Một số trường bây giờ định hướng thành tích với nhiều giáo viên đặt '
  'áp lực lên học sinh. [Học sinh] học như các em nên làm nhưng không '
  'thực sự có ý thức hoặc đầu tư vào việc học của mình. Học sinh ngày '
  'nay tôi thấy nhiều em không tốt bằng thế hệ trước."')

P('Một số người tham gia bày tỏ lo ngại về tác động của thiếu ngủ lên '
  'việc học của học sinh. Một hiệu trưởng ở Hà Nội bày tỏ lo ngại về '
  'mối liên kết giữa thiếu ngủ và điều tiết cảm xúc và học tập kém. Bà '
  'nêu: "Vấn đề lớn nhất của học sinh là không ngủ đủ giấc. Và ở độ tuổi '
  'này, khi [các em] không ngủ đủ giấc, nó sẽ dẫn đến uể oải. Các em hầu '
  'như luôn ở trạng thái uể oải và buồn ngủ. Các em không tập trung. Và '
  'khi trạng thái này kéo dài một thời gian dài, nó dẫn đến những phản '
  'ứng rất tiêu cực. Các em trở nên dễ giận dữ, nóng nảy và thậm chí '
  'kém tập trung hơn vào việc học. Và khi các em không tập trung, kết '
  'quả học tập của các em không cao."')

H('6.5. Các yếu tố liên quan đến gia đình (Family-related Factors)', level=2)

H('6.5.1. Các yếu tố nguy cơ gia đình đối với sức khoẻ tâm thần và hạnh '
  'phúc kém (Family Risk Factors for Poor Mental Health and Well-Being)',
  level=3)

P('Người cung cấp thông tin chủ chốt bày tỏ lo ngại về các yếu tố nguy '
  'cơ khác đối với sức khoẻ tâm thần học sinh kém, bao gồm các yếu tố '
  'gia đình. Các yếu tố nguy cơ dựa trên gia đình bao gồm stress kinh '
  'tế gia đình, bỏ bê của phụ huynh, xung đột gia đình và thiếu kỹ năng '
  'nuôi dạy con.')

P('Người tham gia quan sát rằng các vấn đề kinh tế gia đình đặt học sinh '
  'vào nguy cơ các vấn đề cảm xúc – xã hội. Các nhà quản lý và chuyên gia '
  'lưu ý rằng stress tài chính trong gia đình có thể gây lo âu cho trẻ '
  'em và tác động đến thành tích học tập của các em. Đói nghèo gia đình '
  'cũng có thể dẫn đến tự trọng thấp của học sinh. Một nhà quản lý Sở '
  'GD&ĐT từ Điện Biên nêu: "bạn có thể thấy cách học sinh muốn hoà nhập '
  'vào cuộc sống xã hội bị ngăn cản bởi thiếu hỗ trợ tài chính. Ví dụ, '
  'nếu các em muốn một thứ gì đó phải mất một ít tiền để mua, cuối cùng '
  'đẩy các em trở nên ngại hơn. Các em cố gắng co mình lại."')

P('Một số hiệu trưởng và nhà quản lý xác định bỏ bê của phụ huynh là nguy '
  'cơ tiềm năng đối với các vấn đề sức khoẻ tâm thần vị thành niên. Họ '
  'bày tỏ lo ngại rằng phụ huynh thường phải ưu tiên công việc hoặc kinh '
  'doanh hơn là dành thời gian cho con vị thành niên. Một nhà quản lý '
  'Sở GD&ĐT lưu ý rằng phụ huynh thường tập trung vào kinh doanh và '
  'những thứ khác, và không phải lúc nào cũng coi việc lắng nghe và đồng '
  'cảm với con cái là quan trọng.')

H('6.5.2. Kiến thức sức khoẻ tâm thần và hỗ trợ của phụ huynh '
  '(Parental Mental Health Knowledge and Support)', level=3)

P('Giáo viên và các nhà quản lý từ một số tỉnh lưu ý tầm quan trọng của '
  'sự tham gia của phụ huynh trong việc hỗ trợ học sinh có các mối quan '
  'ngại sức khoẻ tâm thần.')

# --- Trang 53 ---
PageMark('--- Trang 53, UNICEF Việt Nam, 2022 ---')
P('Như một giáo viên ở Hà Nội nêu: "Tôi thấy ngay cả đối với những học '
  'sinh không có bất kỳ vấn đề sức khoẻ tâm thần lớn nào, sự phối hợp '
  'và quan tâm của phụ huynh đối với tâm lý của con là điều quan trọng '
  'nhất. Nếu phụ huynh quan tâm, ngay khi có dấu hiệu thay đổi rất nhỏ '
  'trong tâm lý của con, can thiệp sớm sẽ hiệu quả ngay lập tức. Nhưng '
  'nếu gia đình không hiểu vấn đề thì nó trở thành một vấn đề lớn."')

P('Nhiều người tham gia lo ngại về việc phụ huynh thiếu kiến thức và kỹ '
  'năng về sức khoẻ tâm thần và hạnh phúc vị thành niên. Một người tham '
  'gia từ Sở LĐ-TB&XH Hà Nội bày tỏ lo ngại về kiến thức sức khoẻ tâm '
  'thần của phụ huynh nêu: "Vấn đề là phụ huynh thiếu kỹ năng để có thể '
  'đồng hành với con, để khám phá các vấn đề của con, để dành thời gian '
  'với con. Áp lực cuộc sống và vấn đề kiếm tiền là một vấn đề, nhưng '
  'cũng trong những gia đình rất giàu có, không có nhiều thời gian cho '
  'con cái. [Phụ huynh cần] hiểu tại sao con đang có những thay đổi tâm '
  'lý để có thể hiểu và giúp các em. Để làm tất cả những điều đó, phụ '
  'huynh phải có kiến thức."')

H('6.6. Các yếu tố nguy cơ liên quan đến công nghệ (Technology-Related Risk Factors)',
  level=2)

P('Nhiều người tham gia bày tỏ lo ngại về lượng thời gian học sinh dành '
  'cho sử dụng công nghệ, bao gồm thời gian dành cho lướt internet, chơi '
  'game và mạng xã hội. Các mối quan ngại tồn tại về cả tác động của '
  'việc vị thành niên tiếp xúc với nội dung tiêu cực hoặc không lành '
  'mạnh và tác động của lượng thời gian dành cho sử dụng công nghệ thay '
  'vì tương tác thế giới thực với bạn bè và gia đình. Một nhà quản lý '
  'Sở GD&ĐT nhận xét rằng các vấn đề sức khoẻ tâm thần vị thành niên '
  '"bị tăng cường bởi tác động của tiến bộ công nghệ, mặt tối của các '
  'thiết bị như điện thoại thông minh, trò chơi điện tử hoặc mạng xã '
  'hội. Tất cả những điều này có thể ảnh hưởng nghiêm trọng đến sức '
  'khoẻ tâm thần của học sinh."')

H('6.7. COVID-19, sức khoẻ tâm thần và hạnh phúc '
  '(COVID-19, Mental Health and Well-Being)', level=2)

P('Tất cả các bên liên quan bày tỏ lo ngại về tác động của đại dịch '
  'COVID-19 lên sức khoẻ tâm thần và hạnh phúc học sinh. Các lo ngại '
  'bao gồm tác động của các hạn chế xã hội và cô lập lên sự phát triển '
  'tâm lý – xã hội học sinh và tác động của giáo dục trực tuyến lên sức '
  'khoẻ tâm thần học sinh. Học sinh báo cáo lo ngại về cô lập xã hội '
  'trong COVID-19. Một học sinh ở Hà Nội chia sẻ: "Khi ở nhà, số người '
  'em giao tiếp bị hạn chế hơn, vì không có nhiều người cùng tuổi với '
  'em ở đó. Kiểu như cảm xúc của em không được hiểu tốt nhất."')

P('Giáo viên và các nhà quản lý trường học cũng lo ngại về tác động của '
  'đại dịch lên học sinh. Các mối quan ngại bao gồm cô lập xã hội, '
  'thiếu tiếp cận các hoạt động phát triển lành mạnh bình thường và '
  'tác động của đại dịch lên lịch thi và stress học sinh kết quả. Một '
  'giáo viên từ Hà Nội báo cáo: "Học trực tuyến có nghĩa là trong nhiều '
  'trường hợp bạn ở nhà hầu hết thời gian, điều này ngăn bạn đi ra ngoài '
  'nhiều và do đó gây nhiều stress hơn, được làm tệ hơn bởi khối lượng '
  'công việc tăng do tác động của COVID."')

P('Có mối quan ngại cụ thể dành cho học sinh cuối cấp khi các em chuẩn '
  'bị cho kỳ thi và tuyển sinh đại học. Một chuyên gia tại Bộ GD&ĐT '
  'chia sẻ: "do ảnh hưởng của dịch, học sinh cuối cấp rất lo lắng khi '
  'làm bài thi tốt nghiệp trung học phổ thông hoặc bài thi tốt nghiệp, '
  'vì tình hình dịch bệnh đã ảnh hưởng nghiêm trọng đến các em. Phải '
  'có can thiệp kịp thời."')

H('6.8. Các yếu tố nguy cơ của học sinh dân tộc thiểu số '
  '(Ethnic Minority Student Risk Factors)', level=2)

P('Nhiều người tham gia lo ngại về các yếu tố nguy cơ dựa trên trường '
  'học đối với các vấn đề sức khoẻ tâm thần của học sinh dân tộc thiểu '
  'số. Các yếu tố kinh tế xã hội gia đình được công nhận là các yếu tố '
  'nguy cơ then chốt đối với sức khoẻ tâm thần kém. Tại Gia Lai, một '
  'nhà quản lý lưu ý rằng ly hôn của phụ huynh và bạo lực gia đình tác '
  'động tiêu cực đến hạnh phúc của học sinh. Người này cũng lưu ý rằng '
  'học sinh dân tộc thiểu số thường bị áp lực phải kết hôn ở tuổi vị '
  'thành niên và ở nhà để chăm sóc bố mẹ.')

# --- Trang 54 ---
PageMark('--- Trang 54, UNICEF Việt Nam, 2022 ---')
P('Nhà quản lý này lưu ý rằng ngay cả những học sinh dân tộc thiểu số '
  'rất thành công cũng trở nên tuyệt vọng về tương lai vì các em bị áp '
  'lực phải kết hôn và ở lại làng của mình, nơi cơ hội việc làm ở nhà '
  'rất khan hiếm. "Các em tận mắt nhìn thấy rằng sau giáo dục, không có '
  'cơ hội việc làm. [Các em có thể đi] đại học và thạc sĩ, nhưng khi '
  'các em về nhà các em ở nhà." Việc học của học sinh dân tộc thiểu số '
  'cũng bị tác động bởi một số sự kiện, bao gồm nhiệm vụ nông nghiệp '
  'theo mùa và đám tang. Các nhà quản lý ở Gia Lai bày tỏ lo ngại rằng '
  'vắng mặt trường học thường xuyên làm giảm gắn kết học sinh và đặt '
  'học sinh vào nguy cơ bỏ học.')

P('Một nhà quản lý từ Sở GD&ĐT Gia Lai báo cáo rằng các rào cản sức '
  'khoẻ tâm thần học sinh dân tộc thiểu số bao gồm các rào cản ngôn ngữ '
  'đối với nhiều học sinh và gia đình. Tuy nhiên, một nhà quản lý ở '
  'Điện Biên thấy thiếu kỹ năng cảm xúc – xã hội là cơ bản hơn: "Một '
  'khi các em đã đạt lớp 5, 6, tiếng Việt của các em sẽ khá ổn. Vấn đề '
  'còn lại là rào cản kỹ năng; các kỹ năng mềm về hành vi, giải quyết '
  'vấn đề và kỹ năng giao tiếp không đủ."')

P('Một vấn đề được giáo viên và nhà quản lý trường học bày tỏ là khó '
  'khăn trong việc thu hút phụ huynh dân tộc thiểu số vào các mối quan '
  'ngại về sức khoẻ tâm thần. Như một nhà quản lý ở Sở GD&ĐT Điện Biên '
  'giải thích: "[Một mối quan ngại] là phụ huynh chủ yếu từ cộng đồng '
  'thiểu số và do đó không thể hiểu vấn đề này. Đó là trường hợp chúng '
  'tôi có và chúng tôi có thể thấy cách các vấn đề sức khoẻ tâm thần '
  'của các dân tộc thiểu số hoặc những người ở các khu vực nông thôn đã '
  'thiếu sự chú ý."')

P('Một lo ngại khác về các yếu tố nguy cơ học sinh dân tộc thiểu số là '
  'cô lập địa lý. Nhà quản lý Sở GD&ĐT từ Điện Biên giải thích: "[Đối '
  'với] các nhóm thiểu số, vấn đề đáng chú ý nhất chúng ta có thể thấy '
  'là cô lập. Bất cứ khi nào một học sinh học xa nhà, sẽ khó hơn để các '
  'em thích nghi và bạn có thể thấy càng cố gắng thích nghi, các em '
  'càng trở nên ngại hơn và thiếu tự tin."')

H('6.9. Các yếu tố nguy cơ của học sinh LGBTQ (LGBTQ Student Risk Factors)',
  level=2)

P('Học sinh LGBTQ báo cáo các trải nghiệm khác nhau với môi trường '
  'trường học và các yếu tố nguy cơ liên quan. Nhìn chung, học sinh '
  'LGBTQ chia sẻ các nguồn gây stress chung với học sinh không LGBTQ ở '
  'chỗ các em lo ngại về khối lượng công việc nặng, các kỳ thi sắp tới '
  'và các nguồn stress học tập khác. Ngoài ra, học sinh LGBTQ chia sẻ '
  'các nguồn gây stress về phân biệt đối xử và bắt nạt liên quan đến '
  'tình trạng LGBTQ của mình. Một vài học sinh trong FGD theo học các '
  'trường rất hỗ trợ học sinh LGBTQ. Những học sinh này rất tích cực '
  'về trường mình, với một học sinh chia sẻ: "Vì em học ở [trường], '
  'mọi người có quan điểm rất cởi mở về các vấn đề liên quan đến LGBT, '
  'và các giảng viên quan tâm rất nhiều đến học sinh LGBT. Hầu hết học '
  'sinh ở trường không phán xét về bản dạng giới của chúng em. Các bạn '
  'nói chuyện với chúng em như những người bình thường." Học sinh trong '
  'các môi trường trường học hỗ trợ chia sẻ các ví dụ về giáo viên quan '
  'tâm đến học sinh và dành thời gian để hỗ trợ các em về mặt cảm xúc. '
  'Các học sinh khác báo cáo trải qua phân biệt đối xử ở trường. Một '
  'học sinh cảm thấy không được tôn trọng và bị hiểu lầm khi một giáo '
  'viên nói với cả lớp rằng LGBTQ là một "xu hướng."')

doc.add_page_break()

# ============================================================
# CHUONG 7 — CHINH SACH VA CHUONG TRINH MHPSS HOC DUONG
# ============================================================
H('CHƯƠNG 7: CHÍNH SÁCH VÀ CHƯƠNG TRÌNH MHPSS HỌC ĐƯỜNG CHO HỌC SINH VỊ THÀNH '
  'NIÊN TẠI VIỆT NAM', level=1)
P('(Chapter 7 — School-Based Adolescent Student MHPSS Policies and Programmes in Viet Nam)',
  italic=True, align='center', size=11)

# --- Trang 56 ---
PageMark('--- Trang 56, UNICEF Việt Nam, 2022 ---')

H('Các phát hiện chính (Key Findings)', level=2)

P('1. Bắt đầu từ năm 2005, Bộ GD&ĐT đã thiết lập một số chính sách hỗ trợ sức '
  'khoẻ tâm thần học sinh. Các chính sách này chỉ đạo phát triển các chương trình '
  'tư vấn học đường, cung cấp giáo dục hoà nhập cho trẻ em khuyết tật, giải '
  'quyết tác động tiêu cực của đại dịch COVID-19 lên sức khoẻ tâm thần và hạnh '
  'phúc học sinh, và gần đây nhất, thúc đẩy nhận thức sức khoẻ tâm thần và các '
  'kỹ năng sức khoẻ tâm thần qua Chương trình Sức khoẻ Học đường toàn diện '
  '(2021–2025).', bold=True)

P('2. Học sinh vị thành niên tin rằng có nhu cầu về hỗ trợ sức khoẻ tâm thần '
  'dựa trên trường học, bao gồm tư vấn chuyên nghiệp và các hoạt động thúc đẩy '
  'sức khoẻ tâm thần chung. Các em ngại tìm kiếm trợ giúp từ giáo viên tư vấn '
  'do lo ngại về tính bảo mật và sự không thoải mái liên quan đến việc chia sẻ '
  'thông tin cá nhân với giáo viên.', bold=True)

P('3. Giáo viên, nhà quản lý và chuyên gia đã xác định các thách thức hiện tại '
  'trong việc cung cấp hỗ trợ sức khoẻ tâm thần cho học sinh vị thành niên ở '
  'trường:', bold=True)
P('  a. Các chương trình tư vấn học đường thường có nhân sự là giáo viên có '
  'đào tạo hạn chế và thời gian hạn chế cho hỗ trợ tư vấn.')
P('  b. Trường học thường thiếu một không gian riêng tư dành riêng cho tư vấn '
  'học sinh.')
P('  c. Các bên liên quan xác định các mối quan hệ giáo viên – học sinh mạnh '
  'mẽ, hỗ trợ là thiết yếu để nhận diện sớm và hỗ trợ các vấn đề sức khoẻ tâm '
  'thần học sinh. Tuy nhiên, giáo viên, hiệu trưởng và nhà quản lý thấy nhiều '
  'rào cản đối với các mối quan hệ giáo viên – học sinh hỗ trợ bao gồm sĩ số '
  'lớp lớn, thiếu chú ý đến vấn đề trong các chương trình đào tạo giáo viên, '
  'và thiếu ưu tiên trong chương trình giáo dục.')
P('  d. Chương trình giảng dạy trung học cơ sở và trung học phổ thông tập '
  'trung vào các lớp học học thuật khắt khe và cạnh tranh. Học sinh có ít cơ '
  'hội cho kỹ năng sống, nghệ thuật và thể thao.')

P('4. Liên quan đến khả năng của giáo viên hỗ trợ sức khoẻ tâm thần và hạnh '
  'phúc học sinh là khả năng của họ trong việc tạo động lực hiệu quả cho học '
  'sinh và quản lý hành vi không đúng mực trong lớp học. Trong quá khứ, hình '
  'phạt thể xác được sử dụng để tạo động lực và kỷ luật học sinh. Giáo viên '
  'vẫn có thể tin vào kỷ luật thể chất hoặc có thể không biết các chiến lược '
  'khác để tạo động lực cho học sinh và quản lý hành vi không đúng mực của '
  'học sinh.', bold=True)

P('5. Hầu hết các hiệu trưởng và nhà quản lý Sở GD&ĐT khuyến nghị tạo ra một '
  'vị trí cụ thể cho một tư vấn viên học đường chuyên nghiệp.', bold=True)

P('6. Nhiều người tham gia đã gợi ý nhu cầu về một cách tiếp cận toàn diện để '
  'hỗ trợ sức khoẻ tâm thần và hạnh phúc của học sinh vị thành niên, bao gồm '
  'các tư vấn viên chuyên nghiệp, các chương trình hiểu biết sức khoẻ tâm '
  'thần cho học sinh, giáo viên và phụ huynh, các hệ thống sàng lọc và đánh '
  'giá học sinh, cải thiện các mối quan hệ giáo viên – học sinh và cung cấp '
  'các lớp kỹ năng sống và hạnh phúc cho học sinh.', bold=True)

P('7. Có nhiều rào cản đối với việc có được điều trị chuyên nghiệp cho học '
  'sinh có các vấn đề sức khoẻ tâm thần nghiêm trọng, bao gồm thiếu nguồn lực '
  'cộng đồng và tác động của kỳ thị lên thái độ của phụ huynh đối với việc '
  'tìm kiếm trợ giúp.', bold=True)

P('8. Học sinh dân tộc thiểu số vị thành niên có nhu cầu sức khoẻ tâm thần và '
  'hạnh phúc cụ thể chưa được đáp ứng đầy đủ trong các trường học. Các học '
  'sinh này thường trải qua đói nghèo, phân biệt đối xử và bị gạt ra ngoài lề '
  'xã hội liên quan đến nguy cơ bỏ học và sức khoẻ tâm thần kém. Các tư vấn '
  'viên học đường đối mặt với các thách thức trong việc giao tiếp và thu hút '
  'các gia đình dân tộc thiểu số.', bold=True)

P('9. Học sinh LGBTQ bày tỏ nhu cầu về dịch vụ tư vấn chuyên nghiệp ở trường '
  'và muốn thấy các chương trình giáo dục giới tính toàn diện của trường bao '
  'gồm kiến thức LGBTQ để giảm kỳ thị và phân biệt đối xử đối với học sinh '
  'LGBTQ.', bold=True)

# --- Trang 57 ---
PageMark('--- Trang 57, UNICEF Việt Nam, 2022 ---')

H('7.1. Các chính sách của Bộ GD&ĐT (MOET Policies)', level=2)

P('Năm 2005, Bộ GD&ĐT đã ban hành Công văn số 9971/BGDĐT-HSSV Triển khai Tư '
  'vấn Học đường cho Học sinh Phổ thông và Sinh viên Đại học. Công văn này '
  'đã đưa ra hướng dẫn cho việc phát triển tư vấn học đường và thúc đẩy việc '
  'phát triển ban đầu đào tạo trong tư vấn học đường. Các trường công lập ở '
  'các thành phố đô thị lớn đã khởi động các chương trình tư vấn học đường. '
  'Tuy nhiên, việc phát triển chương trình được phát triển trên cơ sở từng '
  'trường riêng lẻ mà không có nghiên cứu đánh giá hiệu quả chương trình và '
  'không có tiêu chuẩn cho việc phát triển chương trình. Hơn nữa, số lượng '
  'trường học phát triển các chương trình tư vấn học đường như vậy bị hạn '
  'chế.')

P('Năm 2017, Bộ GD&ĐT đã ban hành Thông tư số 31/2017/TT-BGDĐT, một văn bản '
  'pháp lý thiết lập tư vấn học đường trong tất cả các trường công lập tại '
  'Việt Nam. Chính sách này phác thảo một chương trình tư vấn học đường có '
  'chức năng hỗ trợ hạnh phúc tâm thần của học sinh với trọng tâm vào tuổi '
  'học sinh, giới, hôn nhân, gia đình, mang thai vị thành niên, đạo đức, sức '
  'khoẻ, các mối quan hệ, kỹ năng học tập và phát triển nghề nghiệp. Chính '
  'sách cũng nêu rằng các giáo viên có kinh nghiệm với kỹ năng tư vấn học '
  'đường có thể đảm nhận vai trò tư vấn viên học đường và cung cấp tư vấn '
  'sức khoẻ tâm thần cho học sinh, ngoài nhiệm vụ giảng dạy chính của họ. Các '
  'giáo viên đảm nhận vai trò tư vấn viên học đường phải được cung cấp phát '
  'triển chuyên môn có liên quan bởi Sở GD&ĐT, các Trường Đại học Sư phạm và '
  'các chuyên gia tâm lý. Hơn nữa, Thông tư 31 đưa ra các hướng dẫn chính '
  'sách cho việc thiết lập các cơ sở tư vấn học đường trong mọi trường học '
  'tại Việt Nam.')

P('Các chính sách liên quan của Bộ GD&ĐT giải quyết giáo dục cho học sinh '
  'khuyết tật. Quyết định 338 từ tháng 1/2018 ban hành Kế hoạch Giáo dục '
  'toàn diện cho Người Khuyết tật giai đoạn 2018–2020 để thúc đẩy các mục '
  'tiêu giáo dục hoà nhập. Quyết định 1438 tiếp theo vào tháng 10/2018 và '
  'ban hành các mục tiêu cho giai đoạn 2020–2025 để tiếp tục thúc đẩy bảo '
  'vệ, chăm sóc và giáo dục trẻ em khuyết tật trên toàn quốc. Mặc dù các '
  'chính sách này có thể cung cấp cho trẻ em mắc bệnh tâm thần nghiêm trọng, '
  'chúng chủ yếu tập trung vào các khuyết tật thể chất và sẽ không liên '
  'quan đến hầu hết trẻ em có các vấn đề sức khoẻ tâm thần.')

P('Chỉ thị số 31/CT-TTg của Thủ tướng Chính phủ ngày 4/12/2019 thiết lập '
  'các chính sách tăng cường giáo dục đạo đức và lối sống cho học sinh, '
  'sinh viên. Mặc dù không liên quan trực tiếp đến sức khoẻ tâm thần, việc '
  'phát triển giáo dục lối sống có thể bao gồm các khía cạnh của giáo dục '
  'kỹ năng sống sức khoẻ tâm thần, cung cấp các chiến lược phòng ngừa các '
  'vấn đề sức khoẻ tâm thần phổ quát quan trọng.')

P('Chính sách của Bộ GD&ĐT cũng đã hỗ trợ sức khoẻ tâm thần của học sinh bị '
  'tác động bởi đại dịch COVID-19. Chỉ thị số 800/CT-BGDĐT ngày 24/8/2021 '
  'của Bộ trưởng Bộ GD&ĐT đã chỉ đạo thực hiện các nhiệm vụ cho năm học '
  '2021–2022. Chính sách này chỉ đạo ngành Giáo dục quan tâm cẩn thận đến '
  'sức khoẻ tâm thần và hạnh phúc học sinh khi xem xét tác động của đại '
  'dịch, và phát triển hỗ trợ cho học sinh cần. Việc xem xét thêm nhu cầu '
  'sức khoẻ tâm thần và hạnh phúc của học sinh đã được đề cập trong Công '
  'văn số 136/CĐ-BGDĐT ngày 8/2/2022 về hỗ trợ học sinh khi các em trở lại '
  'trường sau một thời gian dài học trực tuyến do đại dịch.')

P('Một chính sách quan trọng về hợp tác liên ngành để hỗ trợ sức khoẻ tâm '
  'thần và hạnh phúc học sinh là Quyết định số 4969/QĐ-BGDĐT ngày 30/12/'
  '2021. Chính sách này ban hành kế hoạch phát triển công tác xã hội trong '
  'ngành giáo dục trong giai đoạn 2021–2025.')

P('Một Chương trình Sức khoẻ Học đường toàn diện cho giai đoạn 2021–2025 đã '
  'được Bộ Giáo dục và Đào tạo phát triển (trình Thủ tướng Chính phủ Quyết '
  'định số 1660/QĐ-TTg ngày 2/10/2021). Nội dung của chương trình tập trung '
  'vào chăm sóc, bảo vệ và quản lý sức khoẻ học sinh, bao gồm đánh giá các '
  'yếu tố nguy cơ đối với sức khoẻ tâm thần, tăng cường truyền thông và '
  'giáo dục sức khoẻ, cải thiện kiến thức của học sinh về sức khoẻ tâm '
  'thần, và cung cấp tư vấn phù hợp với sự phát triển về sinh lý tâm thần '
  'và sức khoẻ tâm thần trong các trường học. Quyết định số 4659/QĐ-BGDĐT '
  'ngày 14/12/2021 của Bộ trưởng Bộ Giáo dục và Đào tạo đã ban hành kế '
  'hoạch thực hiện Quyết định số 1660/QĐ-TTg.')

# --- Trang 58 ---
PageMark('--- Trang 58, UNICEF Việt Nam, 2022 ---')
P('ngày 2/10/2021 của Thủ tướng Chính phủ ban hành Chương trình Sức khoẻ '
  'Học đường giai đoạn 2021–2025. Thủ tướng Chính phủ đã phê duyệt Chương '
  'trình Sức khoẻ Học đường giai đoạn 2021–2025 tại Quyết định số 85/QĐ-'
  'TTg ngày 17/1/2022. Một khía cạnh quan trọng của Chương trình này là kế '
  'hoạch liên kết các chương trình sức khoẻ học đường với các dịch vụ chăm '
  'sóc sức khoẻ cộng đồng để cải thiện chất lượng và hiệu quả của các hoạt '
  'động chăm sóc sức khoẻ học sinh. Một chuyên gia Bộ GD&ĐT chia sẻ rằng '
  'Chương trình này sẽ thúc đẩy nhận thức sức khoẻ tâm thần ở học sinh và '
  'dạy học sinh các chiến lược để cải thiện sức khoẻ tâm thần và hạnh '
  'phúc.')

H('7.2. Nhận thức về các Dịch vụ và Chương trình MHPSS Dựa trên Trường học '
  '(tính sẵn có và hiệu quả)', level=2)

H('7.2.1. Nhận thức của học sinh (Student Perceptions)', level=3)

P('Học sinh đã thảo luận về cách trường của các em hỗ trợ sức khoẻ tâm '
  'thần vị thành niên. Học sinh thường chia sẻ về hỗ trợ giáo viên không '
  'chính thức được cung cấp ở trường. Một học sinh từ Gia Lai báo cáo: '
  '"Giáo viên của em sẽ, ví dụ, đưa lời khuyên cho chúng em nếu chúng em '
  'làm điều gì đó không phù hợp và khuyến khích chúng em thể hiện tốt hơn '
  'để giảm bớt nỗi buồn."')

P('Học sinh cũng đã thảo luận về các dịch vụ tư vấn được cung cấp bởi '
  'trường của các em. Học sinh ở Đồng Tháp nhận thức được một nhóm tư vấn '
  'học đường bao gồm giáo viên và các nhà quản lý ở trường của các em. Tuy '
  'nhiên, các học sinh gặp khó khăn trong việc nhớ những giáo viên nào '
  'trong nhóm. Khi các học sinh ở Điện Biên được hỏi liệu trường của các '
  'em có bất kỳ nguồn lực nào không, các em nói: "Có. Cô [tên giáo viên] '
  'và một giáo viên khác. Cô bảo chúng em nói chuyện với cô nếu có chuyện '
  'gì xảy ra. Nhưng không ai làm." Học sinh ở Hà Nội không đồng ý với nhau '
  'về các dịch vụ tư vấn học đường. Khi được hỏi liệu các em có thể sử '
  'dụng phòng tư vấn trong tương lai nếu các em cần hỗ trợ, các học sinh '
  'nhìn chung đồng ý rằng các em có thể sẽ không sử dụng. Các lý do bao '
  'gồm sợ rằng những gì các em chia sẻ sẽ được chia sẻ với người khác và '
  'vì nó sẽ "xấu hổ."')

H('7.2.2. Nhận thức của giáo viên (Teacher Perceptions)', level=3)

P('Giáo viên ở Đồng Tháp mô tả một dịch vụ tư vấn liên quan đến các giáo '
  'viên chủ nhiệm và hiệu trưởng cung cấp các buổi tư vấn hàng tuần. Các '
  'giáo viên này cũng nhận ra rằng các mối quan hệ giáo viên – học sinh là '
  'then chốt, và các giáo viên phải biết mỗi trẻ em và điều chỉnh hướng '
  'dẫn theo điểm mạnh và nhu cầu của mỗi trẻ em.')

P('Giáo viên ở một trường bán công tại Hà Nội báo cáo rằng trường của họ '
  'có một phòng tư vấn tâm lý có nhân sự là một giáo viên với một số đào '
  'tạo tư vấn. Họ chia sẻ rằng học sinh sử dụng các dịch vụ tư vấn và tin '
  'rằng giáo viên tư vấn và trường được cống hiến để hỗ trợ học sinh về '
  'mặt tâm lý. Một giáo viên báo cáo: "Năm nay, trường sẽ thuê một nhân '
  'viên toàn thời gian chuyên về công tác sức khoẻ tâm thần, thay vì bán '
  'thời gian. Tôi nghĩ hy vọng nó sẽ hiệu quả hơn."')

P('Dữ liệu định lượng của giáo viên về các nguồn lực và chương trình của '
  'trường cho sức khoẻ tâm thần và hạnh phúc học sinh phát hiện rằng 81 % '
  'giáo viên báo cáo kiến thức về các chính sách và/hoặc các chương trình '
  'cơ bản cho sức khoẻ tâm thần học sinh, bao gồm các cơ sở tư vấn học '
  'đường và nhân viên y tế trường học. Khi được hỏi liệu các nguồn lực '
  'và chương trình hiện có có đầy đủ để hỗ trợ học sinh không, 58 % giáo '
  'viên nói rằng không. Xem chi tiết trong Phụ lục 3.')

H('7.2.3. Nhận thức của Hiệu trưởng, Nhà quản lý và Chuyên gia Bộ ngành '
  '(Principals, Administrators and Ministry Experts\' Perceptions)', level=3)

P('Các cuộc phỏng vấn hiệu trưởng đã làm rõ rằng các chương trình và dịch '
  'vụ được cung cấp không nhất quán trên khắp cả nước. Một hiệu trưởng ở '
  'Gia Lai không quen thuộc với các chính sách của Bộ GD&ĐT về sức khoẻ '
  'tâm thần học sinh. Khi được hỏi về hỗ trợ hoặc các chương trình cho '
  'học sinh với nhu cầu cảm xúc – xã hội, người này đã chia sẻ về các '
  'chương trình và hoạt động ngoại khoá được cung cấp tại trường.')

# --- Trang 59 ---
PageMark('--- Trang 59, UNICEF Việt Nam, 2022 ---')
P('Một hiệu trưởng trường ở Hà Nội tin rằng trường của bà cung cấp hỗ trợ '
  'cảm xúc – xã hội hiệu quả cho học sinh. Hiệu trưởng này báo cáo về '
  'phòng và dịch vụ tư vấn học đường, nhưng tập trung vào cách cấu trúc '
  'lớp và chương trình giảng dạy của trường hỗ trợ học sinh bằng cách ưu '
  'tiên các mối quan hệ giáo viên – học sinh và một văn hoá trường học hỗ '
  'trợ. Các giáo viên chủ nhiệm có hai tiết học với học sinh mỗi ngày nên '
  'họ có các mối quan hệ mạnh hơn với học sinh và có thể quan sát học '
  'sinh kỹ và nhận ra các dấu hiệu của các vấn đề sớm.')

P('Một hiệu trưởng khác ở Hà Nội đã xác định ba thành phần trong cách '
  'tiếp cận của trường để hỗ trợ sức khoẻ tâm thần và hạnh phúc của học '
  'sinh vị thành niên, bao gồm một phòng và dịch vụ tư vấn, các mối quan '
  'hệ giáo viên với học sinh và gia đình cho phép họ cung cấp hỗ trợ tư '
  'vấn không chính thức cho học sinh đang vật lộn, và những thay đổi '
  'trong chương trình giảng dạy được thiết kế để cung cấp thêm hỗ trợ '
  'tâm lý cho học sinh.')

P('Các nhà quản lý Sở GD&ĐT đồng ý rằng một khía cạnh cơ bản của hỗ trợ '
  'học sinh dựa trên trường học là các mối quan hệ giữa giáo viên và học '
  'sinh. Học sinh phải có giáo viên đáng tin cậy mà các em có thể đến '
  'nếu các em cần hỗ trợ.')

P('Một chuyên gia từ Bộ GD&ĐT bày tỏ tự tin vào các chương trình sức '
  'khoẻ tâm thần và hạnh phúc hiện tại, và vào hỗ trợ sức khoẻ tâm thần '
  'có sẵn cho học sinh vị thành niên. Chuyên gia chia sẻ rằng đã có nhiều '
  'chương trình bao gồm "tư vấn về sức khoẻ và sinh lý, ngoài đào tạo về '
  'học tập của học sinh. Về chủ đề tổ chức và thêm các hoạt động để hỗ '
  'trợ sức khoẻ tâm thần cho học sinh, có những hoạt động rất thành '
  'công liên quan đến các câu lạc bộ, ví dụ các câu lạc bộ liên quan '
  'đến thể thao hoặc câu lạc bộ sức khoẻ hoặc các hoạt động ngoại khoá '
  'hoặc các cuộc trò chuyện trực tiếp. Giáo dục kỹ năng sống cũng có '
  'thể giúp trẻ em đầu tiên phát triển tài năng và sở thích của mình."')

H('7.3. Nhận thức của Giáo viên và Lãnh đạo Trường về Năng lực Trường và '
  'Giáo viên trong việc cung cấp Hỗ trợ Tâm lý – Xã hội cho Học sinh',
  level=2)

P('Giáo viên bày tỏ lo ngại và nhạy cảm đối với học sinh cần hỗ trợ tâm '
  'lý – xã hội. Họ chia sẻ các ví dụ về những tình huống khi họ bối rối '
  'về cách hỗ trợ học sinh tốt nhất. Ví dụ, một giáo viên ở Đồng Tháp '
  'không chắc chắn cách giúp một học sinh đang vật lộn: "Trong lớp của '
  'tôi, có một học sinh hành động như thể cô ấy bị tự kỷ. Học sinh đó '
  'hầu như không bao giờ nói câu hoàn chỉnh, và đôi khi khi cô ấy nhìn '
  'vào một bài toán hoặc một đoạn văn trên bảng trong lớp, cô ấy hiểu '
  'nó nhưng luôn đánh vần và viết sai. Đôi khi tôi đứng cạnh cô ấy và '
  'bảo cô ấy viết lại các câu, và cô ấy lại mắc lỗi tương tự. Khi tôi '
  'liên hệ với phụ huynh để hợp tác, họ nói rằng cô ấy chỉ như vậy ở '
  'nhà." Sự bối rối của giáo viên về bản chất các khó khăn của học '
  'sinh và cách hỗ trợ tốt nhất cho cô ấy phản ánh sự thiếu năng lực.')

P('Các giáo viên khác báo cáo thiếu kỹ năng để tạo động lực cho học '
  'sinh. Một giáo viên chia sẻ về cách giáo viên trước đây có thể đánh '
  'học sinh để tạo động lực cho các em và bây giờ điều đó không được '
  'phép. Giáo viên này báo cáo khó khăn trong việc tìm cách mới để tạo '
  'động lực cho học sinh và quản lý lớp học.')

P('Các phòng tư vấn học đường thường được nhân sự bởi một giáo viên đã '
  'nhận được một số đào tạo cơ bản trong tư vấn học sinh. Một rào cản '
  'chính đối với dịch vụ tư vấn hiệu quả tại các trường liên quan đến '
  'vấn đề thời gian của giáo viên – tư vấn viên. Giáo viên tại một '
  'trường ở Hà Nội lưu ý các hạn chế của các dịch vụ tư vấn được cung '
  'cấp: "Tôi nghĩ dịch vụ tư vấn tại trường của tôi đã tồn tại 1 đến 2 '
  'năm. Học sinh đã được cung cấp hỗ trợ tâm thần cơ bản nhưng thời '
  'gian cho các nhiệm vụ tư vấn khá hạn chế do giáo viên tư vấn phải '
  'dạy ở nhiều lớp."')

# --- Trang 60 ---
PageMark('--- Trang 60, UNICEF Việt Nam, 2022 ---')

P('Các nhà quản lý cũng bày tỏ lo ngại về năng lực của giáo viên trong '
  'việc cung cấp hỗ trợ sức khoẻ tâm thần. Một nhà quản lý Sở GD&ĐT từ '
  'Gia Lai mô tả các vấn đề với các nỗ lực của giáo viên cung cấp hỗ trợ '
  'không chính thức, bao gồm khó khăn trong việc xác định chính xác học '
  'sinh cần, và thu hút gia đình không nhất quán. Chuyên gia này cũng '
  'thấy các vấn đề liên quan đến dịch vụ tư vấn hiện tại được cung cấp, '
  'bao gồm việc chương trình tư vấn không có nhân sự phù hợp và nhất '
  'quán do giáo viên được giao nhiệm vụ tư vấn khi họ có khoảng trống '
  'trong lịch trình. Điều này dẫn đến sự không nhất quán trong nhóm tư '
  'vấn, làm suy yếu niềm tin và các mối quan hệ của học sinh với các '
  'tư vấn viên. Một vấn đề then chốt khác được xác định là các giáo '
  'viên được giao nhiệm vụ tư vấn không nhận được đào tạo phù hợp để '
  'đảm nhận vai trò. Nhà quản lý Sở GD&ĐT Gia Lai đã lưu ý rằng tài '
  'trợ là vấn đề nền tảng vì các trường không được cung cấp tài trợ '
  'cho đào tạo.')

P('Các nhà quản lý Sở GD&ĐT chia sẻ lo ngại về đào tạo giáo viên và '
  'chương trình giảng dạy thiếu hỗ trợ cho sức khoẻ tâm thần và hạnh '
  'phúc của học sinh. Một nhà quản lý Sở GD&ĐT từ Hà Nội chia sẻ: "Tôi '
  'đã học lớp phương pháp giảng dạy ở đại học, nhiều nhất họ dạy chúng '
  'tôi là cách dạy bài học để học sinh hiểu và không nhiều về việc '
  'chăm sóc học sinh."')

P('Một chuyên gia Bộ GD&ĐT chia sẻ các rào cản đối với sức khoẻ tâm '
  'thần dựa trên trường học bao gồm thiếu các cơ sở vật chất (bao '
  'gồm các phòng tư vấn riêng tư) và thiếu nguồn nhân lực được đào '
  'tạo đầy đủ cho tư vấn.')

P('Dữ liệu định lượng của giáo viên chỉ ra rằng hầu hết giáo viên cảm '
  'thấy tự tin về kiến thức cơ bản của mình về sức khoẻ tâm thần học '
  'sinh, mặc dù một nhóm đáng kể (khoảng 36 %) tin rằng họ không có '
  'đủ thông tin về sức khoẻ tâm thần. Giáo viên cũng đồng ý phần lớn '
  'rằng việc xác định và hỗ trợ học sinh có các vấn đề sức khoẻ tâm '
  'thần là một phần vai trò của họ.')

P('Khi được hỏi về đào tạo sức khoẻ tâm thần học sinh vị thành niên, '
  'hầu hết giáo viên (86,4 %) báo cáo rằng họ không nhận được đào tạo '
  'nào. Xem kết quả khảo sát chi tiết trong Phụ lục 3.')

H('7.4. Nhận thức của Học sinh, Giáo viên và Lãnh đạo Trường về Nhu cầu '
  'MHPSS Dựa trên Trường học', level=2)

H('7.4.1. Nhận thức của học sinh về các nhu cầu MHPSS dựa trên trường '
  'học (Student perceptions of school-based MHPSS needs)', level=3)

P('Khi được hỏi về điều gì khác các em muốn từ trường để hỗ trợ hạnh '
  'phúc của các em, học sinh nêu rằng các em rất muốn có một người '
  'được chỉ định nói chuyện tại trường về các vấn đề của mình và muốn '
  'có một không gian riêng tư để nói chuyện về các chủ đề nhạy cảm. '
  'Học sinh nhận ra nhu cầu đánh giá tâm lý ở trường. Một học sinh ở '
  'Gia Lai chia sẻ: "Em biết một bạn trong lớp em cứ ngồi một mình vào '
  'giờ nghỉ, nên em quyết định đi nói chuyện với bạn đó để cổ vũ bạn. '
  'Em nghĩ ban quản lý nhà trường nên kiểm tra tâm lý của bạn."')

P('Một vấn đề quan tâm của thanh thiếu niên là phản ứng của giáo viên '
  'đối với các mối quan hệ và tình dục của học sinh. Một học sinh từ '
  'Gia Lai chia sẻ: "Em nghĩ các giáo viên không nên có quá nhiều vấn '
  'đề với sự phát triển câu chuyện tình yêu giữa các học sinh."')

P('Học sinh lưu ý rằng có các lễ kỷ niệm được tổ chức thường xuyên ở '
  'trường sẽ giúp giảm stress và cải thiện tinh thần của các em.')

P('Học sinh cũng mong trường sẽ giảm áp lực học tập đặt lên học sinh '
  'và coi đây là một vấn đề quan trọng trong việc giảm stress của các '
  'em. Học sinh ở Hà Nội đã gợi ý rằng trường có thể giao ít bài tập '
  'học hơn để giảm khối lượng công việc và giáo viên không nên mắng '
  'học sinh.')

# --- Trang 62 ---
PageMark('--- Trang 62, UNICEF Việt Nam, 2022 ---')
AddImg('p062_img1_Im0.jpg',
       'Hình 5 (trang 62). Học sinh nam — chương Chính sách & Chương trình MHPSS học đường',
       'Schoolboy — School-based MHPSS Programmes chapter')

H('7.4.2. Nhận thức của Giáo viên, Hiệu trưởng, Nhà quản lý về Nhu cầu '
  'MHPSS Dựa trên Trường học', level=3)

P('Kết quả khảo sát định lượng của giáo viên bao gồm nhận thức của '
  'giáo viên về các rào cản chính đối với việc cung cấp thêm các dịch '
  'vụ cho trẻ em có nhu cầu sức khoẻ tâm thần ở trường. Giáo viên xác '
  'định một số rào cản chính, bao gồm thiếu nguồn nhân lực (như các '
  'tư vấn viên), thiếu sự quan tâm/lo ngại từ phụ huynh, thiếu cơ hội '
  'đào tạo cho nhân viên trường và thiếu chính sách để hướng dẫn các '
  'dịch vụ. Xem Phụ lục 3 để biết kết quả chi tiết.')

P('Giảng viên và lãnh đạo trường đã xác định các cách để cải thiện '
  'MHPSS tại các trường. Hầu hết các hiệu trưởng và nhà quản lý Sở '
  'GD&ĐT, bao gồm cả những người từ Gia Lai và Điện Biên, khuyến '
  'nghị tạo ra một vị trí cụ thể cho một tư vấn viên học đường chuyên '
  'nghiệp. Các nhà quản lý trường học từ mọi tỉnh bày tỏ lo ngại về '
  'động lực và năng lực của giáo viên và nhân viên y tế trong việc '
  'đảm nhận vai trò tư vấn.')

P('Một hiệu trưởng từ một trường bán công ở Hà Nội chia sẻ: "Có hai mô '
  'hình tại các trường. Mô hình đầu tiên là sẽ không có tư vấn viên '
  'chuyên nghiệp. Nó hơi hiệu quả nhưng không quá hiệu quả. Bởi vì đó '
  'không phải là chuyên ngành của [giáo viên], họ chỉ làm từ kinh '
  'nghiệm và từ những gì họ quan sát được khi dạy học sinh trung học. '
  'Mô hình thứ hai giống như mô hình ở trường chúng tôi. Chúng tôi '
  'thuê một tư vấn viên được đào tạo để hỗ trợ học sinh. Hiệu quả của '
  'mô hình này cao hơn nhiều so với mô hình đầu tiên."')

P('Trong khi Thông tư 31 đã thúc đẩy tiến bộ lớn trong việc thiết lập '
  'các không gian tư vấn riêng tư trong mỗi trường, các cuộc phỏng '
  'vấn đã làm rõ rằng nhiều trường vẫn thiếu các cơ sở vật chất này. '
  'Một nhà quản lý từ một tỉnh nông thôn ước tính rằng 70 % các trường '
  'trong huyện thiếu một phòng cụ thể, riêng tư cho tư vấn.')

P('Nhiều người tham gia đã gợi ý nhu cầu về một cách tiếp cận toàn '
  'diện để hỗ trợ sức khoẻ tâm thần và hạnh phúc của học sinh vị '
  'thành niên.')

# --- Trang 63 ---
PageMark('--- Trang 63, UNICEF Việt Nam, 2022 ---')
P('Họ lưu ý nhu cầu về các tư vấn viên chuyên nghiệp nhưng cũng đề '
  'cập đến nhu cầu về các chương trình khác. Các ý tưởng được chia '
  'sẻ bao gồm sự tham gia và đào tạo về sức khoẻ tâm thần của học '
  'sinh, giáo viên và phụ huynh, các hệ thống sàng lọc và đánh giá '
  'học sinh, cải thiện các mối quan hệ giáo viên – học sinh và cung '
  'cấp các lớp kỹ năng sống và hạnh phúc cho học sinh.')

P('• Sự tham gia của học sinh vào chương trình sức khoẻ tâm thần. '
  'Các nhà quản lý Sở GD&ĐT bày tỏ lo ngại rằng học sinh thường '
  'thiếu nhận thức về các vấn đề sức khoẻ tâm thần của chính mình '
  'và ngại tìm kiếm trợ giúp. Một rào cản liên quan đến việc học '
  'sinh tìm kiếm trợ giúp là kỳ thị.', bold=True)

P('• Đào tạo giáo viên về sức khoẻ tâm thần. Một số người tham gia '
  'KI báo cáo nhu cầu về đào tạo giáo viên chung về sức khoẻ tâm '
  'thần vị thành niên. Như một chuyên gia từ Sở GD&ĐT Hà Nội nêu: '
  '"Mong muốn của tôi là có một khoá đào tạo cho giáo viên đi sâu '
  'vào các triệu chứng sức khoẻ tâm thần, phát hiện sớm, và cách '
  'phát hiện các dấu hiệu tinh tế của các vấn đề sức khoẻ tâm thần '
  'hoặc các sự kiện sang chấn."', bold=True)

P('• Đào tạo và thu hút phụ huynh. Giáo viên và nhà quản lý cũng '
  'lưu ý nhu cầu về cải thiện giao tiếp và hợp tác với phụ huynh '
  'trong nỗ lực này.', bold=True)

P('• Mối quan hệ giáo viên – học sinh. Các nhà quản lý Sở Y tế cũng '
  'lưu ý rằng các vấn đề này thường bị bỏ qua bởi phụ huynh và giáo '
  'viên.', bold=True)

P('Đối với học sinh có các vấn đề sức khoẻ tâm thần nghiêm trọng cần '
  'chăm sóc chuyên gia, những người tham gia thảo luận nhu cầu phát '
  'triển các hệ thống giới thiệu cho hỗ trợ bên ngoài, và vượt qua '
  'các rào cản gia đình và tài trợ trong những tình huống này.')

P('Ngoài chương trình tư vấn, nhiều giáo viên và nhà quản lý lưu ý '
  'tầm quan trọng của việc sửa đổi chương trình giảng dạy hoặc văn '
  'hoá giáo dục để hỗ trợ tốt hơn hạnh phúc học sinh vị thành niên. '
  'Ở Hà Nội, các đề xuất của giáo viên bao gồm (a) giảm áp lực học '
  'tập, (b) tăng các khoá/cơ hội nghệ thuật và thể thao cho học '
  'sinh, và (c) dạy thêm kỹ năng sống.')

# --- Trang 64 ---
PageMark('--- Trang 64, UNICEF Việt Nam, 2022 ---')

H('Khoảng trống tài trợ (Funding Gaps)', level=3)

P('Trong khi một đánh giá đầy đủ về cơ cấu tài trợ sức khoẻ học sinh '
  'vị thành niên vượt quá phạm vi nghiên cứu này, những người tham '
  'gia thường bày tỏ lo ngại về vấn đề tài trợ. Một nhà quản lý Sở '
  'GD&ĐT từ Gia Lai báo cáo rằng, trong khi Thông tư 31 phát triển '
  'các dịch vụ tư vấn học đường, nó không cung cấp tài trợ cho nỗ '
  'lực này. Nhà quản lý này mô tả cách tài trợ phải đến từ các ngân '
  'sách trường học đã chật hẹp, rất khó khăn, đặc biệt xem xét tình '
  'trạng thiếu giáo viên hiện tại.')

H('Tiếp cận các nguồn lực sức khoẻ tâm thần cộng đồng '
  '(Access to Community-Based Mental Health Resources)', level=3)

P('Một lĩnh vực khác có nhu cầu về hỗ trợ cho sức khoẻ tâm thần và '
  'hỗ trợ tâm lý – xã hội của học sinh là tính sẵn có và khả năng '
  'tiếp cận các nguồn lực sức khoẻ tâm thần dựa trên cộng đồng. Học '
  'sinh không nhận thức được bất kỳ nguồn lực địa phương nào cho sức '
  'khoẻ tâm thần. Học sinh ở Điện Biên nhận thức được một bệnh viện '
  'tâm thần gần đó, nhưng khi được hỏi ai nhận được trợ giúp ở đó, '
  'các em nói đó là dành cho "người điên," người có "tâm trí không '
  'ổn định," và những người sử dụng ma tuý.')

P('Những người tham gia khác cũng lưu ý sự thiếu hụt các dịch vụ dựa '
  'trên cộng đồng để đáp ứng nhu cầu sức khoẻ tâm thần của trẻ em. '
  'Trong Điện Biên, một trung tâm bảo trợ xã hội cấp tỉnh đã được xây '
  'dựng nhưng được báo cáo là bị sử dụng dưới mức đáng kể do thiếu '
  'nhân viên và nguồn nhân lực.')

P('Kỳ thị là một rào cản khác đối với các gia đình cần chăm sóc tâm '
  'lý chuyên nghiệp cho con. Nhiều phụ huynh ngại tìm kiếm chăm sóc '
  'tâm thần cho con vì sợ những hậu quả xã hội tiêu cực nếu cộng đồng '
  'biết về vấn đề của đứa trẻ.')

H('Thách thức về các chương trình sức khoẻ tâm thần cho học sinh '
  'Dân tộc Thiểu số và LGBTQ', level=3)

P('Các thách thức đặc biệt trong việc giải quyết nhu cầu sức khoẻ '
  'tâm thần và hạnh phúc của học sinh dân tộc thiểu số đã được chia '
  'sẻ bởi một nhà quản lý Sở GD&ĐT.')

# --- Trang 65 ---
PageMark('--- Trang 65, UNICEF Việt Nam, 2022 ---')
P('từ Điện Biên đã báo cáo rằng các tư vấn viên học đường gặp khó '
  'khăn trong việc thu hút và phối hợp với phụ huynh học sinh dân '
  'tộc thiểu số. Người này lưu ý: "Các tư vấn viên học đường chỉ có '
  'thể hợp tác giữa các phụ huynh ở khu vực đô thị. Trong thời đại '
  'công nghệ và thông tin, chúng ta có thể giao tiếp qua các nền tảng '
  'mạng xã hội như Zalo, Facebook, điều mà tôi nghĩ khu vực đô thị '
  'đã tận dụng khá tốt. Nhưng trong so sánh, đối với các dân tộc '
  'thiểu số ở khu vực nông thôn có một hạn chế đặt ra bởi bối cảnh '
  'địa lý và từ sự hiểu biết của họ về [tiếng Việt] và mức độ quen '
  'thuộc với công nghệ thông tin."')

P('Những người tham gia bày tỏ nhu cầu đặc biệt về các chương trình '
  'sức khoẻ tâm thần cho học sinh dân tộc thiểu số. Nhà quản lý ở '
  'Điện Biên nêu: "Tôi nghĩ [phát triển chương trình sức khoẻ tâm '
  'thần] là một vấn đề rất quan trọng và thiết yếu cho sự phát triển '
  'của ngành giáo dục tại nước ta, đặc biệt là đối với nhóm thiểu '
  'số. Điều này là vì những trẻ em chúng ta đang làm việc chỉ là '
  'những học sinh trẻ nhưng phải chịu nhiều tác động và bất lợi về '
  'phát triển kinh tế – lịch sử và bố cục địa lý."')

P('Học sinh LGBTQ báo cáo các trải nghiệm khác nhau với hỗ trợ tâm '
  'lý – cảm xúc trường học. Hỗ trợ không chính thức dưới hình thức '
  'giáo viên quan tâm được học sinh trải nghiệm là rất tích cực. Tuy '
  'nhiên, học sinh LGBTQ trong FGD không thấy các hỗ trợ tâm lý '
  'chính thức là hữu ích. Một học sinh báo cáo rằng trường của các '
  'em có một phòng tâm lý nhưng không tin rằng học sinh đến đó để '
  'được giúp đỡ. Khi được hỏi tại sao không, học sinh suy đoán: "Em '
  'nghĩ có lẽ là, đến đó sẽ không làm được gì, hoặc các em sợ phải '
  'đối mặt với giáo viên vì các giáo viên trong phòng tâm lý cũng là '
  'giáo viên trong trường. Vì vậy, em cảm thấy cô ấy có thể biết tất '
  'cả mọi thứ em đã nói với cô ấy và em cảm thấy không thoải mái." '
  'Khi được hỏi trường còn có thể làm gì khác để hỗ trợ sức khoẻ '
  'tâm thần và hạnh phúc của học sinh LGBTQ, học sinh trong FGD đồng '
  'ý rằng các trường nên cung cấp các chương trình giáo dục giới '
  'tính chính thức hơn, bao gồm giáo dục về xu hướng tính dục và '
  'đa dạng giới.')

doc.add_page_break()

# ============================================================
# CHUONG 8 — MOLISA + MOH + HOP TAC LIEN NGANH
# ============================================================
H('CHƯƠNG 8: CHÍNH SÁCH VÀ CHƯƠNG TRÌNH MHPSS CHO HỌC SINH VỊ THÀNH NIÊN — KHU '
  'VỰC XÃ HỘI VÀ Y TẾ, HỢP TÁC LIÊN NGÀNH', level=1)
P('(Chapter 8 — Adolescent Student MHPSS Policies and Programmes in Viet Nam: Social '
  'and Health Sectors, and Cross-Sector Collaboration)', italic=True, align='center', size=11)

# --- Trang 67 ---
PageMark('--- Trang 67, UNICEF Việt Nam, 2022 ---')

H('Các phát hiện chính (Key Findings)', level=2)

P('1. Từ năm 2011, Bộ LĐ-TB&XH đã thiết lập các chính sách cung cấp cho việc phát '
  'triển chăm sóc sức khoẻ tâm thần dựa trên cộng đồng cho trẻ em và người lớn.',
  bold=True)

P('2. Các chuyên gia và nhà quản lý Bộ LĐ-TB&XH báo cáo nhu cầu về một chương '
  'trình sức khoẻ tâm thần cụ thể cho trẻ em và vị thành niên tập trung vào '
  'phòng ngừa các vấn đề sức khoẻ tâm thần thông qua thúc đẩy an toàn trong gia '
  'đình và thông qua dạy phụ huynh, vị thành niên và thành viên cộng đồng cách '
  'nhận biết các dấu hiệu sớm của các vấn đề sức khoẻ tâm thần và nơi tìm trợ '
  'giúp.', bold=True)

P('3. Các lĩnh vực nhu cầu được xác định cho Bộ LĐ-TB&XH về sức khoẻ tâm thần '
  'học sinh vị thành niên bao gồm:', bold=True)
P('  a. Thiết lập các tiêu chuẩn cho chất lượng dịch vụ và các thủ tục M&E hiệu '
  'quả để cải thiện liên tục các dịch vụ cộng đồng.')
P('  b. Phát triển nguồn nhân lực sức khoẻ tâm thần cộng đồng.')

P('4. Từ năm 1999, Bộ Y tế đã chịu trách nhiệm về chăm sóc sức khoẻ tâm thần dựa '
  'trên cộng đồng. Trọng tâm ban đầu là các bệnh tâm thần nặng ở người lớn. Từ '
  'năm 2005, Bộ Y tế đã mở rộng trọng tâm sang các rối loạn sức khoẻ tâm thần '
  'phổ biến, bao gồm trầm cảm và lo âu, và sức khoẻ tâm thần trẻ em.', bold=True)

P('5. Bộ Y tế hiện đang phát triển các chính sách và chương trình tập trung vào '
  'các vấn đề sức khoẻ tâm thần phổ biến của trẻ em và vị thành niên, bao gồm '
  'rối loạn lo âu, trầm cảm và rối loạn hành vi.', bold=True)
P('  a. Các lĩnh vực nhu cầu được xác định cho Bộ Y tế về sức khoẻ tâm thần học '
  'sinh vị thành niên bao gồm:')
P('  b. Đào tạo cho các nhà quản lý cấp huyện và nhân viên y tế trường học về '
  'sức khoẻ tâm thần vị thành niên.')
P('  c. Phát triển nguồn nhân lực, cụ thể trong lĩnh vực tâm thần vị thành niên.')

P('6. Hợp tác giữa Giáo dục và Khu vực Xã hội sẽ được cải thiện bởi:', bold=True)
P('  a. Các hệ thống chia sẻ thông tin về học sinh có nguy cơ và cần hỗ trợ.')
P('  b. Hỗ trợ của khu vực xã hội cho học sinh vị thành niên qua các chương '
  'trình hiểu biết sức khoẻ tâm thần trong trường học và tư vấn cho học sinh.')

P('7. Hợp tác giữa Giáo dục và Khu vực Y tế sẽ được cải thiện bởi:', bold=True)
P('  a. Đào tạo cho nhân viên y tế trường học trong các lĩnh vực sàng lọc sức '
  'khoẻ tâm thần dựa trên trường học, đánh giá nguy cơ và giao thức chăm sóc '
  'sức khoẻ tâm thần để giới thiệu học sinh cần chăm sóc sức khoẻ tâm thần '
  'chuyên nghiệp.')
P('  b. Hỗ trợ khu vực y tế cho việc cải thiện hiểu biết sức khoẻ tâm thần giữa '
  'học sinh, giáo viên và phụ huynh để tạo thuận lợi cho việc nhận diện sớm và '
  'can thiệp các vấn đề sức khoẻ tâm thần học sinh.')

P('8. Các chuyên gia đồng ý rằng cần có một Chiến lược Quốc gia về Sức khoẻ Tâm '
  'thần liên quan đến hợp tác và điều phối các khu vực y tế, giáo dục và xã hội.',
  bold=True)

# --- Trang 68 ---
PageMark('--- Trang 68, UNICEF Việt Nam, 2022 ---')

H('8.1. Các chính sách của Bộ LĐ-TB&XH (MOLISA Policies)', level=2)

P('Quyết định số 1215/QĐ-TTg của Thủ tướng Chính phủ từ tháng 7/2011 cung cấp hỗ '
  'trợ xã hội và phục hồi dựa trên cộng đồng cho người mắc bệnh tâm thần và rối '
  'loạn tâm thần giai đoạn 2011–2020. Chính sách này đặt ra các mục tiêu phòng '
  'ngừa và chăm sóc các rối loạn tâm thần trong cộng đồng. Các hoạt động được '
  'phác thảo bởi chính sách bao gồm (a) phát triển các cơ sở Trung tâm Bảo trợ '
  'Xã hội, bao gồm 20 cơ sở trên khắp cả nước và ít nhất 3 Trung tâm Bảo trợ Xã '
  'hội khu vực, (b) tăng năng lực cấp tỉnh để hỗ trợ người mắc rối loạn tâm '
  'thần, (c) tăng nguồn nhân lực cho phòng ngừa và điều trị các rối loạn tâm '
  'thần qua đào tạo cho nhân viên xã hội và các gia đình bị ảnh hưởng, và (d) '
  'thí điểm và mở rộng quy mô các mô hình phòng ngừa và điều trị rối loạn tâm '
  'thần dựa trên cộng đồng hiệu quả. Thông tư thiết lập hợp tác liên ngành trong '
  'nỗ lực này. Bộ Y tế được chỉ đạo hỗ trợ xây dựng năng lực điều trị sức khoẻ '
  'tâm thần bằng cách tích hợp nghiên cứu và chuyên môn vào đào tạo và phát '
  'triển. Bộ GD&ĐT được chỉ đạo tích hợp giáo dục kỹ năng sống cho học sinh vào '
  'các trường học, thực hiện các chương trình trong các trường để phòng ngừa '
  'và can thiệp sớm cho học sinh có các rối loạn tâm thần, và phối hợp với '
  'ngành y tế để chăm sóc học sinh có các rối loạn tâm thần.')

P('Quyết định này được nối tiếp bởi Quyết định số 1929/QĐ-TTg của Thủ tướng '
  'Chính phủ ngày 25/11/2020 phê duyệt chương trình hỗ trợ xã hội và phục hồi '
  'tâm thần dựa trên cộng đồng, bao gồm cho trẻ em tự kỷ và người có rối loạn '
  'tâm thần, giai đoạn 2021–2030. Quyết định số 1069/QĐ-LĐTBXH ngày 27/9/2021 '
  'của Bộ Lao động, Thương binh và Xã hội đã thiết lập Kế hoạch triển khai '
  'Chương trình về hỗ trợ xã hội và phục hồi tâm thần, trẻ em tự kỷ và rối '
  'loạn tâm thần dựa trên cộng đồng giai đoạn 2021–2030.')

P('Quyết định số 112/QĐ-TTg của Thủ tướng Chính phủ tháng 1/2021 đã thiết lập '
  'chính sách tiếp tục thúc đẩy phát triển công tác xã hội trong khung thời '
  'gian 2021–2030 bằng cách nâng cao nhận thức công chúng về các hoạt động '
  'công tác xã hội và cải thiện chất lượng các dịch vụ công tác xã hội. Trong '
  'khi Quyết định này tập trung vào phát triển công tác xã hội rộng hơn, sức '
  'khoẻ tâm thần được đề cập cụ thể trong mục 2.7.a, chỉ đạo Bộ LĐ-TB&XH '
  '"truyền thông và nâng cao nhận thức của các cấp chính quyền, các ngành và '
  'cộng đồng về vai trò và vị trí của công tác xã hội và các cơ sở cung cấp '
  'dịch vụ công tác xã hội trong chăm sóc sức khoẻ tâm thần và chăm sóc sức '
  'khoẻ cho người khuyết tật, người cao tuổi, trẻ em có hoàn cảnh đặc biệt và '
  'các nhóm thiệt thòi khác."')

P('Đáng quan trọng cụ thể đối với phát triển chăm sóc sức khoẻ tâm thần dựa '
  'trên trường học là Phần IV, Điều 2.5 giải quyết vai trò của Bộ GD&ĐT: "Bộ '
  'Giáo dục và Đào tạo (MOET) chịu trách nhiệm chính và phối hợp với các Bộ, '
  'ngành và cơ quan liên quan để… cải thiện chất lượng đào tạo viên công tác '
  'xã hội và thiết lập mạng lưới công chức và nhân viên công tác xã hội ở các '
  'trường học." Ngoài ra, các mục tiêu bao gồm mở rộng nhân viên công tác xã '
  'hội trong các trường, với ít nhất 50 % cơ sở giáo dục cung cấp dịch vụ '
  'công tác xã hội như được lập kế hoạch đến năm 2025 (Mục tiêu 2.a), và đạt '
  'ít nhất 60 % cơ sở giáo dục đến năm 2030 so với năm 2025 (Mục tiêu 2.b).')

P('Các chính sách liên quan bao gồm Quyết định số 1437/QĐ-TTg của Thủ tướng '
  'Chính phủ (được phê duyệt tháng 10/2018) về phê duyệt đề án 2018–2025 về '
  'chăm sóc phát triển toàn diện trẻ em trong những năm đầu đời của gia đình '
  'và đời sống cộng đồng. Chính sách này chứng minh nỗ lực quốc gia để thúc '
  'đẩy phát triển ấu thơ, đảm bảo phát triển thể chất, nhận thức, cảm xúc và '
  'tiếp cận bình đẳng với các dịch vụ hỗ trợ cho trẻ em. Tuy nhiên, vì nó tập '
  'trung vào trẻ em dưới tám tuổi, nó không liên quan trực tiếp đến sức khoẻ '
  'tâm thần vị thành niên.')

# --- Trang 69 ---
PageMark('--- Trang 69, UNICEF Việt Nam, 2022 ---')

P('Luật Trẻ em năm 2016 cung cấp nền tảng pháp lý cho các quyền của trẻ em tại '
  'Việt Nam và thể chế hoá các hướng dẫn và chính sách đảm bảo việc thực thi '
  'các quyền trẻ em theo tinh thần Công ước Liên Hợp Quốc về Quyền Trẻ em. '
  'Luật Trẻ em cung cấp các quy định đảm bảo nhiều quyền trẻ em, bao gồm quyền '
  'riêng tư, quyền sống với cha mẹ, quyền được nhận nuôi và quyền được bảo vệ '
  'khỏi lạm dụng bao gồm bạo lực, quấy rối tình dục, bóc lột lao động, bỏ rơi '
  'và bắt cóc. Mặc dù Luật này là chính sách quan trọng liên quan đến giảm các '
  'yếu tố nguy cơ chính đối với các vấn đề sức khoẻ tâm thần trẻ em, nó không '
  'cụ thể giải quyết hoặc cung cấp hỗ trợ cho sức khoẻ tâm thần và hạnh phúc '
  'vị thành niên.')

P('Một chuyên gia từ Bộ LĐ-TB&XH đã chia sẻ nhiều cách chung mà các chính '
  'sách này hỗ trợ sức khoẻ tâm thần và hạnh phúc trẻ em. Tuy nhiên, chuyên '
  'gia này cho rằng Bộ LĐ-TB&XH có thể cần phát triển các chính sách và '
  'chương trình sức khoẻ tâm thần trẻ em cụ thể hơn. Như bà nêu: "Thành thật '
  'mà nói, các chương trình và dự án hiện tại có sẵn, nhưng chúng tôi có thể '
  'cần ban hành một chương trình riêng, một dự án sức khoẻ tâm thần riêng '
  'cho trẻ em. Cần có các giải pháp can thiệp cụ thể để phòng ngừa các rối '
  'loạn tâm thần ở trẻ em, cũng như cung cấp hỗ trợ kịp thời cho trẻ em."')

P('Theo chuyên gia này, các chương trình phòng ngừa nên (1) thúc đẩy môi '
  'trường gia đình an toàn để phòng ngừa sang chấn, (2) đưa cho phụ huynh các '
  'công cụ để nhận biết các dấu hiệu sớm của các vấn đề sức khoẻ tâm thần, '
  '(3) dạy trẻ em về sức khoẻ tâm thần và các dấu hiệu sớm của các vấn đề và '
  'cách giảm thiểu rủi ro, và (4) cung cấp các bộ công cụ cho hiểu biết sức '
  'khoẻ tâm thần cơ bản rộng rãi, và để cải thiện năng lực cộng đồng để phát '
  'hiện các dấu hiệu cảnh báo sớm của các vấn đề và thực hiện các giới thiệu '
  'phù hợp.')

H('8.2. Các chương trình và dịch vụ của Khu vực Xã hội '
  '(Social Sector Programmes and Services)', level=2)

P('Khu vực dịch vụ của Bộ LĐ-TB&XH bao gồm một số dịch vụ dựa trên trường học. '
  'Như chuyên gia Bộ LĐ-TB&XH chia sẻ, một số trường đã kết hợp các nhà cung '
  'cấp công tác xã hội để cung cấp hỗ trợ tâm lý cho học sinh. Tuy nhiên, được '
  'lưu ý rằng ở hầu hết các trường, nhân viên y tế phụ trách sức khoẻ của trẻ '
  'em, nhưng họ thường là bán thời gian và không được trang bị để bao gồm sức '
  'khoẻ tâm thần. Khu vực dịch vụ xã hội cũng có các cán bộ quản lý trẻ em '
  'trong cộng đồng làm việc với một nhóm cộng tác viên công tác xã hội trong '
  'quản lý trường hợp.')

P('Bộ LĐ-TB&XH cũng hỗ trợ đào tạo cho các gia đình, học sinh và thành viên '
  'cộng đồng trong các lĩnh vực liên quan đến sức khoẻ và hạnh phúc trẻ em. '
  'Ví dụ, một nhà quản lý Sở LĐ-TB&XH Hà Nội đã chia sẻ về một chương trình '
  'đào tạo cho phụ huynh về kỹ năng nuôi dạy con, bao gồm khả năng của phụ '
  'huynh nhận biết và hỗ trợ các vấn đề của trẻ em.')

P('Sở LĐ-TB&XH Hà Nội cung cấp các dịch vụ thông qua Trung tâm Công tác Xã '
  'hội Hà Nội. Mặc dù các dịch vụ này chủ yếu tập trung vào trẻ em nhỏ tuổi, '
  'một người cung cấp thông tin chủ chốt từ Sở LĐ-TB&XH Hà Nội đã báo cáo '
  'rằng mục tiêu của trung tâm công tác xã hội là "hỗ trợ trẻ em và gia đình '
  'trong các trường hợp gia đình gặp khó khăn trong nuôi dạy con hoặc các '
  'vấn đề của trẻ em."')

# --- Trang 70 ---
PageMark('--- Trang 70, UNICEF Việt Nam, 2022 ---')
AddImg('p070_img1_Im0.jpg',
       'Hình 6 (trang 70). Học sinh nữ tươi cười — chương MHPSS Khu vực Y tế & Xã hội',
       'Smiling schoolgirl — Health & Social Sectors chapter')
P('Một dịch vụ khác của Bộ LĐ-TB&XH được cung cấp từ năm 2017 là đường dây '
  'nóng bảo vệ trẻ em quốc gia 111 của Cục Trẻ em. Đường dây nóng này cung '
  'cấp hỗ trợ điện thoại 24 giờ cho nạn nhân vị thành niên của lạm dụng, '
  'khai thác và buôn bán người. Cục Trẻ em có ba trung tâm tổng đài ở các '
  'vùng Bắc, Trung và Nam của đất nước để giải quyết sự khác biệt địa lý '
  'trong giọng nói để giúp trẻ em gọi từ tất cả các khu vực cảm thấy thoải '
  'mái. Các báo cáo từ Cục Trẻ em thuộc Bộ LĐ-TB&XH cho thấy sau gần 16 '
  'năm vận hành đường dây nóng 111, Cục đã nhận hơn 4 triệu cuộc gọi để '
  'thảo luận các vấn đề liên quan đến trẻ em. Theo nhà quản lý Bộ LĐ-TB&XH, '
  'nhân viên đường dây nóng đều được đào tạo về các vấn đề sức khoẻ tâm '
  'thần ở trẻ em.')

P('Nguồn nhân lực được báo cáo là một mối quan ngại đối với việc cung cấp '
  'dịch vụ khu vực xã hội. Một chuyên gia Bộ LĐ-TB&XH mô tả sự thiếu hụt '
  'đào tạo trong sức khoẻ tâm thần trẻ em: "Ở cấp bộ có một số người chuyên '
  'về sức khoẻ tâm thần. Cục Trẻ em có một Phòng Chăm sóc Trẻ em có nhân '
  'viên được giao trực tiếp giám sát các vấn đề liên quan đến sức khoẻ tâm '
  'thần. Nhưng khi bạn chuyển sang cấp tỉnh, huyện và xã thì không có '
  'chuyên gia. Trong sở LĐ-TB&XH, bạn sẽ có một người phụ trách tất cả các '
  'vấn đề trẻ em, bao gồm cả sức khoẻ tâm thần."')

# --- Trang 71 ---
PageMark('--- Trang 71, UNICEF Việt Nam, 2022 ---')
P('Các nhân viên xã hội ở "cấp cơ sở đó gánh một khối lượng công việc rất '
  'bận rộn, tôi nghĩ đó cũng là một rào cản lớn, một vấn đề rất khó khăn." '
  'Suy nghĩ một cách toàn diện, một người tham gia từ Sở LĐ-TB&XH Hà Nội đã '
  'xem xét cách cộng đồng có thể hỗ trợ hạnh phúc trẻ em và vị thành niên '
  'tích cực bằng cách cung cấp thêm cơ hội cho sự tham gia vào các hoạt '
  'động tích cực: "Trong ngành giáo dục, trong các tổ chức và đoàn thể, '
  'chẳng hạn như Hội Phụ nữ và Đoàn Thanh niên, họ có thể làm nhiều điều '
  'liên quan đến trẻ em, bao gồm cải thiện tương tác cha mẹ – con, cải '
  'thiện môi trường trường học và tạo ra các đội, nhóm hoạt động, câu lạc '
  'bộ cho trẻ em."')

H('8.3. Các chính sách của Bộ Y tế (MOH Policies)', level=2)

P('Hỗ trợ của Bộ Y tế cho các rối loạn tâm thần bắt đầu từ năm 1999 với '
  'Quyết định của Thủ tướng Chính phủ về Dự án Chăm sóc Sức khoẻ Tâm thần '
  'Dựa trên Cộng đồng, giao phó điều trị sức khoẻ tâm thần cho ngành y tế. '
  'Quyết định này nhấn mạnh việc thực hiện một chương trình sức khoẻ tâm '
  'thần cộng đồng. Năm 2000, Bộ trưởng Bộ Y tế đã giao nhiệm vụ cho Dịch '
  'vụ Sức khoẻ Tâm thần khởi động Chương trình Mục tiêu Quốc gia về Sức '
  'khoẻ Tâm thần, cho toàn bộ dân số. Từ 2000 đến 2006, Chương trình Chăm '
  'sóc Sức khoẻ Tâm thần nhắm vào chăm sóc và điều trị tâm thần phân liệt. '
  'Trọng tâm sau đó được mở rộng để bao gồm một chương trình cho động kinh. '
  'Các hoạt động của chương trình này bao gồm đào tạo sức khoẻ tâm thần '
  'cho nhân viên y tế và các cộng tác viên y tế, cũng như các khảo sát hộ '
  'gia đình để xác định bệnh nhân trầm cảm và động kinh, giao thuốc hàng '
  'tháng cho bệnh nhân, và giám sát và hỗ trợ bệnh nhân thông qua thuốc '
  'và giáo dục sức khoẻ qua phương tiện truyền thông thôn.')

P('Từ năm 2015, chương trình đã bắt đầu tập trung vào các vấn đề sức khoẻ '
  'tâm thần phổ biến, bao gồm trầm cảm và lo âu, và sức khoẻ tâm thần trẻ '
  'em, bao gồm tự kỷ và ADHD. Các rào cản chính đối với tiến bộ trong lĩnh '
  'vực này bao gồm thiếu nhân viên y tế được đào tạo và có kinh nghiệm '
  'trong sức khoẻ tâm thần trẻ em và vị thành niên.')

P('Các chính sách gần đây liên quan đến sức khoẻ tâm thần bao gồm Quyết '
  'định số 155/QĐ-TTg của Thủ tướng Chính phủ ngày 29/1/2022 mà qua đó '
  'chính phủ đã phê duyệt Kế hoạch Quốc gia về phòng chống các bệnh không '
  'lây nhiễm và các rối loạn sức khoẻ tâm thần giai đoạn 2022–2025. Quyết '
  'định số 712/QĐ-BYT ngày 21/3/2022 do Bộ Y tế ban hành thiết lập một kế '
  'hoạch phát triển công tác xã hội trong ngành y tế giai đoạn 2021–2030.')

P('Theo một chuyên gia Bộ Y tế, Bộ Y tế hiện không có chính sách cụ thể '
  'về sức khoẻ tâm thần của học sinh. Như chuyên gia mô tả, chương trình '
  'chăm sóc sức khoẻ tâm thần cộng đồng tập trung vào động kinh, cung cấp '
  'thuốc miễn phí và điều trị cho trẻ em mắc động kinh. Chuyên gia lưu ý '
  'rằng: "Tuy nhiên, các rối loạn tâm thần ở trẻ em do phát triển sinh '
  'thái – xã hội và công nghiệp hoá gây ra thì không được giải quyết, như '
  'các rối loạn cảm xúc, rối loạn hành vi, tự tử ở vị thành niên hoặc trẻ '
  'em. Thành thật mà nói, những vấn đề đó không được giải quyết." Ông chia '
  'sẻ rằng Bộ Y tế, Cục Y tế Dự phòng cùng với Cục Quản lý Khám, chữa '
  'bệnh đang phát triển một kế hoạch hành động bao gồm sức khoẻ tâm thần '
  'trong lĩnh vực bệnh không lây nhiễm.')

H('8.4. Các chương trình và dịch vụ của Khu vực Y tế '
  '(Health Sector Programmes and Services)', level=2)

P('Bộ Y tế và các sở y tế tỉnh dường như không có các chương trình dựa '
  'trên trường học về sức khoẻ tâm thần học sinh. Một nhà quản lý từ Sở '
  'Y tế Hà Nội đã chia sẻ rằng bộ có một chương trình để cải thiện chăm '
  'sóc sức khoẻ sinh sản cho vị thành niên. Chương trình này liên quan '
  'đến việc đến các trường trung học cơ sở.')

# --- Trang 72 ---
PageMark('--- Trang 72, UNICEF Việt Nam, 2022 ---')
P('và trung học phổ thông để nói chuyện với học sinh về sức khoẻ sinh sản, '
  'bao gồm phát triển sinh học và tâm lý. Chương trình cũng bao gồm các '
  'khoá đào tạo giáo viên về chăm sóc sức khoẻ sinh sản vị thành niên để '
  'giáo viên có các kỹ năng hỗ trợ học sinh có vấn đề. Bộ Y tế cũng cung '
  'cấp các khoá đào tạo ngắn cho nhân viên y tế trong lĩnh vực chăm sóc '
  'sức khoẻ sinh sản cho vị thành niên.')

P('Mặc dù nhận thức được tầm quan trọng của sức khoẻ tâm thần vị thành '
  'niên, các chuyên gia y tế đồng ý rằng họ thiếu đào tạo về sức khoẻ '
  'tâm thần và không đủ trang bị để hỗ trợ các trường hoặc học sinh '
  'trong lĩnh vực này. Một nhà quản lý phụ trách phòng y tế học đường '
  'của Đồng Tháp đã chia sẻ rằng bà chưa có đào tạo về tâm thần, hoặc '
  'bất cứ điều gì liên quan đến sức khoẻ tâm thần trẻ em và vị thành '
  'niên.')

P('Các cuộc phỏng vấn với các chuyên gia sở y tế tỉnh chỉ ra nhu cầu đào '
  'tạo cho khu vực y tế công cộng về sức khoẻ tâm thần trẻ em và vị '
  'thành niên, và nhu cầu phát triển các liên kết với các nguồn lực bệnh '
  'viện tâm thần cấp tỉnh.')

P('Một rào cản chính đối với hỗ trợ của khu vực y tế cho sức khoẻ tâm '
  'thần học sinh vị thành niên là thiếu các chuyên gia y tế chuyên sâu. '
  'Theo chuyên gia Bộ Y tế: "Các bác sĩ chuyên về sức khoẻ tâm thần '
  'không đủ, chưa kể các bác sĩ sức khoẻ tâm thần nhi khoa. Có rất ít '
  'người được đào tạo trong chăm sóc sức khoẻ tâm thần nhi khoa." Các '
  'nguyên nhân nền tảng của nguồn nhân lực không đủ bao gồm kỳ thị từ '
  'trong ngành y tế, nơi các bác sĩ tâm thần không được tôn trọng, và '
  'kỳ thị từ trong văn hoá và cộng đồng, nơi sức khoẻ tâm thần không '
  'được coi trọng.')

# --- Trang 73 ---
PageMark('--- Trang 73, UNICEF Việt Nam, 2022 ---')
P('Một rào cản nguồn nhân lực khác liên quan đến việc thiếu một mã nghề '
  'nghiệp chuyên nghiệp cho Nhà tâm lý học. Như người tham gia Bộ Y tế '
  'mô tả, Bộ Y tế đang chờ hoàn thiện một đề xuất về vấn đề này và kế '
  'hoạch sau đó trình đề xuất cho Bộ Nội vụ.')

H('8.5. Hợp tác liên ngành về sức khoẻ tâm thần học sinh vị thành niên '
  '(Cross-Sector Collaboration)', level=2)

H('8.5.1. Hợp tác giữa Giáo dục và Khu vực Xã hội '
  '(Education and Social Sector Collaboration)', level=3)

P('Bằng chứng cho hợp tác liên ngành giữa Bộ GD&ĐT và Bộ LĐ-TB&XH bao gồm '
  'các chương trình hỗ trợ thực hiện Luật Trẻ em 2016 về quyền trẻ em '
  'trong các trường. Chương trình này bao gồm một quy trình chung để hỗ '
  'trợ trẻ em, bao gồm quản lý trường hợp từ Sở LĐ-TB&XH, và điều phối '
  'với các trường và cơ quan y tế xung quanh nhu cầu của từng trẻ em.')

P('Mặc dù những nỗ lực này đáng hứa hẹn, có bằng chứng cho thấy điều '
  'phối và hợp tác giữa các khu vực giáo dục và xã hội còn thiếu, và '
  'phần lớn các nỗ lực điều phối đã tập trung vào trẻ em nhỏ tuổi hơn '
  'thay vì vị thành niên. Điều này được xác nhận bởi một chuyên gia Bộ '
  'LĐ-TB&XH người đã báo cáo rằng điều phối giữa hai ngành vẫn "rất phân '
  'mảnh."')

P('Một lĩnh vực cần cải thiện hợp tác là chia sẻ thông tin và dữ liệu '
  'liên quan. Một nhà quản lý Sở GD&ĐT từ Hà Nội bày tỏ lo ngại về điều '
  'phối giữa các ngành giáo dục và xã hội. Người này hiểu rằng Sở LĐ-'
  'TB&XH có các hoạt động liên quan đến sức khoẻ tâm thần vị thành niên, '
  'nhưng chưa thấy dữ liệu về công việc này.')

P('Một chuyên gia từ Bộ LĐ-TB&XH đã chia sẻ về cơ hội mở rộng hợp tác '
  'MOET-MOLISA cho sức khoẻ tâm thần học sinh vị thành niên. Gần đây, '
  'Bộ GD&ĐT đã làm việc với Bộ LĐ-TB&XH về Dự án Sức khoẻ Học đường '
  'Quốc gia giai đoạn 2021–2025 được Thủ tướng Chính phủ phê duyệt '
  'trong Quyết định số 1660 ngày 2/10/2021 và được khởi động ngày '
  '10/2/2022.')

H('8.5.2. Hợp tác giữa Giáo dục và Khu vực Y tế '
  '(Education and Health Sector Collaboration)', level=3)

P('Các ngành giáo dục và y tế đã có một số hợp tác trong lĩnh vực sức '
  'khoẻ tâm thần học sinh vị thành niên. Một nhà quản lý từ Sở Y tế Hà '
  'Nội đã chia sẻ các ví dụ về hợp tác bao gồm Sở Y tế cung cấp đào tạo '
  'cho giáo viên về.')

# --- Trang 74 ---
PageMark('--- Trang 74, UNICEF Việt Nam, 2022 ---')
P('rất ít điều phối xảy ra. Các nhà quản lý Sở GD&ĐT từ Gia Lai và Điện '
  'Biên đã chia sẻ rằng nhân viên y tế trường học không được đào tạo về '
  'tư vấn và không có kết nối với các bệnh viện tâm thần hoặc các sở '
  'cấp huyện hoặc cấp tỉnh.')

P('Một chuyên gia Bộ Y tế đã khẳng định sự thiếu điều phối chính thức '
  'giữa Bộ Y tế, Bộ GD&ĐT và Bộ LĐ-TB&XH trong lĩnh vực sức khoẻ tâm '
  'thần học sinh. Người này báo cáo rằng ba bộ đã đề xuất các lĩnh vực '
  'hợp tác trong việc thực hiện Chiến lược Quốc gia về Sức khoẻ Tâm '
  'thần nhưng chính phủ đã tái cấu trúc chiến lược để tập trung vào '
  'một chiến lược cho mỗi Bộ. Chuyên gia Bộ Y tế khẳng định rằng Bộ Y '
  'tế muốn hợp tác với Bộ GD&ĐT về sức khoẻ tâm thần vị thành niên '
  'nhưng "vấn đề là không có giao thức để duy trì loại hợp tác đó, để '
  'giữ nó thường xuyên. Không có giao tiếp. Vì vậy, chúng tôi chỉ làm '
  'công việc của mình. Nếu có một giao thức chỉ ra rằng chúng tôi nên '
  'có các cuộc họp thường xuyên, trao đổi thông tin, thì sự hợp tác '
  'giữa hai phòng ban để làm nhiệm vụ A, nhiệm vụ B sẽ diễn ra tự '
  'nhiên."')

P('Các người tham gia giáo dục và y tế bày tỏ sự nhiệt tình với hợp '
  'tác và chia sẻ nhiều ý tưởng về cách hai ngành có thể làm việc cùng '
  'nhau để mang lại lợi ích cho sức khoẻ tâm thần vị thành niên. Một '
  'nhà quản lý từ Sở Y tế Gia Lai đã hăng hái xem xét các cơ hội để '
  'làm việc với các trường và lưu ý rằng ngành y tế có thể nâng cao '
  'nhận thức sức khoẻ tâm thần giữa giáo viên và học sinh và thúc đẩy '
  'các giới thiệu đến bệnh viện cho các vấn đề nghiêm trọng. Một nhà '
  'quản lý từ Sở GD&ĐT Gia Lai chia sẻ sự nhiệt tình này và gợi ý một '
  'cuộc họp hàng năm giữa lãnh đạo trường và các khoa hoặc bệnh viện '
  'tâm thần.')

P('Hợp tác giữa Bộ Y tế và Bộ GD&ĐT sẽ được tối ưu hoá bởi các vai trò '
  'rõ ràng về các nỗ lực sức khoẻ tâm thần học sinh vị thành niên. Các '
  'chuyên gia lưu ý các lĩnh vực cần hợp tác, bao gồm sàng lọc, đánh '
  'giá nguy cơ, giao thức chăm sóc sức khoẻ tâm thần và kế hoạch hành '
  'động.')

P('Chuyên gia Bộ Y tế khẳng định nhu cầu đáng kể về chăm sóc sức khoẻ '
  'tâm thần vị thành niên. Chuyên gia chia sẻ rằng Bộ Y tế nên tập '
  'trung trước vào cải thiện kiến thức sức khoẻ tâm thần của học sinh '
  'vị thành niên và phụ huynh: "Nếu chúng ta chỉ hỗ trợ học sinh mà '
  'không có thay đổi từ môi trường sống hoặc gia đình, sẽ thực sự '
  'khó. Phải có sự đồng bộ hoá, từ môi trường gia đình đến xã hội đến '
  'trường học."')

P('Một người tham gia từ Sở Y tế Hà Nội đã củng cố tầm quan trọng của '
  'đào tạo giáo viên để nhận diện sớm các vấn đề sức khoẻ tâm thần '
  'học sinh. Ông đề xuất rằng: "Giáo viên nên được đào tạo và cung '
  'cấp kiến thức cơ bản về sức khoẻ tâm thần vị thành niên. Họ cần '
  'được trang bị các công cụ để sàng lọc, đánh giá các rối loạn phổ '
  'biến. Ví dụ, họ nên có một danh sách kiểm tra ngắn cho các triệu '
  'chứng. Khi một học sinh có triệu chứng, giáo viên có thể sử dụng '
  'danh sách kiểm tra đó để xác định [học sinh có nguy cơ] và sau đó '
  'mời các bác sĩ tâm thần làm việc với học sinh đó."')

P('Một nhà quản lý từ Đồng Tháp đã thảo luận nhu cầu hợp tác và gợi ý '
  'rằng các ngành y tế, xã hội và giáo dục cùng nhận đào tạo và làm '
  'việc trên một chính sách chung cho hợp tác và cung cấp dịch vụ sức '
  'khoẻ tâm thần. Bà phản ánh rằng điều này đã được thực hiện trong '
  'lĩnh vực dinh dưỡng và nghĩ rằng quy trình cho dinh dưỡng có thể '
  'phục vụ như một mô hình hữu ích để xây dựng hợp tác liên ngành '
  'trong sức khoẻ tâm thần vị thành niên.')

doc.add_page_break()

# ============================================================
# CHUONG 9 — KET LUAN VA KHUYEN NGHI
# ============================================================
H('CHƯƠNG 9: KẾT LUẬN VÀ KHUYẾN NGHỊ (Conclusions and Recommendations)', level=1)

# --- Trang 76-77 KET LUAN ---
PageMark('--- Trang 76–77, UNICEF Việt Nam, 2022 ---')
AddImg('p076_img1_Im0.jpg',
       'Hình 7 (trang 76). Học sinh nữ DTTS trong lớp học — chương Kết luận',
       'Ethnic minority schoolgirls in classroom — Conclusions chapter')

H('KẾT LUẬN (Conclusions)', level=2)

P('UNICEF đã hợp tác với Bộ Giáo dục để nghiên cứu mối quan hệ giữa các yếu tố '
  'liên quan đến trường học và sức khoẻ tâm thần cũng như hạnh phúc của các em '
  'gái và em trai vị thành niên. Kết quả của nghiên cứu này sẽ cung cấp thông '
  'tin cho việc phát triển các can thiệp dựa trên trường học nhắm vào sức khoẻ '
  'tâm thần và hạnh phúc vị thành niên. Nghiên cứu tích hợp dữ liệu từ các khu '
  'vực địa lý khác nhau của Việt Nam, xuyên các ngành liên quan (giáo dục, y '
  'tế và xã hội), và các cấp của hệ thống (học sinh, gia đình, trường, quản '
  'trị địa phương và quốc gia).')

P('Học sinh vị thành niên tại Việt Nam là những người trẻ thông minh với những '
  'ước mơ lớn. Các em yêu học những điều mới, quan tâm đến gia đình và bạn bè, '
  'và phấn đấu để năng động và thành công ở trường và ở nhà. Học sinh vị thành '
  'niên báo cáo một loạt các vấn đề sức khoẻ tâm thần, bao gồm các vấn đề cảm '
  'xúc, các vấn đề bạn bè, tăng động và các vấn đề hành vi. Khoảng 26 % học '
  'sinh báo cáo các triệu chứng hiện tại liên quan đến nguy cơ trung bình hoặc '
  'cao về các vấn đề sức khoẻ tâm thần. Các em gái có nguy cơ cao hơn về các '
  'vấn đề cảm xúc so với các em trai, và học sinh trung học phổ thông có nguy '
  'cơ cao hơn so với học sinh trung học cơ sở.')

P('Tất cả các bên liên quan nhận ra tầm quan trọng của sức khoẻ tâm thần và '
  'hạnh phúc vị thành niên. Giáo viên, hiệu trưởng và các nhà quản lý từ các '
  'ngành giáo dục, xã hội và y tế đều đồng ý rằng sức khoẻ tâm thần vị thành '
  'niên là lĩnh vực đáng quan ngại. Trong khi phụ huynh vật lộn để hiểu và hỗ '
  'trợ con tuổi teen, phụ huynh thể hiện ít kiến thức hơn về các triệu chứng '
  'và yếu tố nguy cơ sức khoẻ tâm thần học sinh, và bày tỏ ít quan ngại hơn '
  'về vấn đề này. Học sinh dân tộc thiểu số và học sinh LGBTQ có thể có nguy '
  'cơ đặc biệt về các vấn đề sức khoẻ tâm thần và hạnh phúc do nhiều yếu tố '
  'nguy cơ khác nhau.')

P('Một số yếu tố liên quan đến trường học đặt học sinh vào nguy cơ các vấn đề '
  'sức khoẻ tâm thần. Các yếu tố khí hậu trường học có liên quan vừa phải với '
  'các vấn đề sức khoẻ tâm thần học sinh, với nhận thức của học sinh về an '
  'toàn trường học, gắn kết học sinh ở trường và môi trường trường học đều có '
  'liên quan đến sức khoẻ tâm thần học sinh. Áp lực học tập có liên quan mạnh '
  'mẽ với sức khoẻ tâm thần học sinh. Các trải nghiệm của học sinh về áp lực '
  'học tập, lo lắng về điểm, chán nản liên quan đến học tập, tự kỳ vọng và '
  'khối lượng công việc đều có liên quan đáng kể với các vấn đề sức khoẻ tâm '
  'thần. Các em gái có nguy cơ cao hơn về áp lực học tập so với các em trai. '
  'Thiếu ngủ của học sinh và các trải nghiệm với bắt nạt đều đáng quan ngại '
  'đối với học sinh và các bên liên quan.')

P('Kết quả phát hiện mối quan ngại rộng rãi giữa các người tham gia về áp lực '
  'học tập. Có sự đồng thuận chung rằng áp lực từ giáo viên và phụ huynh, khối '
  'lượng công việc cao, và stress liên quan đến điểm và kỳ thi có tác động '
  'tiêu cực lên sức khoẻ tâm thần học sinh, gây hại cho phát triển cảm xúc, '
  'xã hội và học tập của học sinh và đặt một số học sinh vào nguy cơ các quỹ '
  'đạo cuộc sống tiêu cực.')

P('Các chính sách và chương trình giáo dục về sức khoẻ tâm thần học sinh vị '
  'thành niên nhằm thiết lập các dịch vụ tư vấn cơ bản trong các trường trung '
  'học cơ sở và trung học phổ thông, bao gồm các cơ sở và nguồn nhân lực được '
  'đào tạo. Mặc dù đã có tiến bộ, bằng chứng chỉ ra các khoảng trống giữa '
  'chính sách và thực hiện. Nhiều trường thiếu các phòng tư vấn được chỉ định '
  'và các tư vấn viên được đào tạo đầy đủ. Hầu hết các trường đặt các giáo '
  'viên hoặc nhà quản lý được đào tạo kém vào vai trò tư vấn viên.')

P('Các chính sách khu vực xã hội trực tiếp cung cấp cho việc phát triển công '
  'tác xã hội cho chăm sóc sức khoẻ tâm thần, bao gồm chăm sóc trong bối cảnh '
  'các trường học. Tuy nhiên, các chương trình cụ thể giải quyết sức khoẻ tâm '
  'thần vị thành niên trong các trường còn thiếu, và hiện tại có rất ít cung '
  'cấp dịch vụ công tác xã hội trực tiếp trong các trường. Hợp tác giữa các '
  'khu vực xã hội và giáo dục phải được tăng cường để mang lại lợi ích cho '
  'học sinh vị thành niên.')

P('Các chính sách khu vực y tế giao phó điều trị sức khoẻ tâm thần cho ngành '
  'y tế, nhưng không trực tiếp giải quyết sức khoẻ tâm thần học sinh vị '
  'thành niên. Các chương trình khu vực y tế trong các trường giải quyết sức '
  'khoẻ tâm thần một cách gián tiếp. Nhân viên y tế trong các trường thiếu '
  'đào tạo và năng lực để hỗ trợ chăm sóc sức khoẻ tâm thần. Và có sự thiếu '
  'điều phối giữa các trường và các nhà cung cấp dịch vụ tâm thần.')

P('Những điểm mạnh của hệ thống bao gồm sự hiểu biết rộng rãi về tầm quan '
  'trọng của sức khoẻ tâm thần học sinh vị thành niên đối với các kết quả '
  'phát triển, và cam kết mạnh mẽ cải thiện các chính sách, chương trình và '
  'dịch vụ được cung cấp để hỗ trợ học sinh.')

P('Mặc dù không được đánh giá trực tiếp trong nghiên cứu hiện tại, có bằng '
  'chứng đáng kể rằng COVID-19 đã gây stress đáng kể cho trẻ em và gia đình '
  'trên khắp thế giới, và tại Việt Nam. Các bên liên quan chia sẻ lo ngại '
  'về tác động của đại dịch lên sức khoẻ và hạnh phúc học sinh.')

# --- Trang 78 KHUYEN NGHI A ---
doc.add_page_break()
PageMark('--- Trang 78, UNICEF Việt Nam, 2022 ---')

H('KHUYẾN NGHỊ (Recommendations)', level=2)

P('Các khuyến nghị sau được đưa ra để Bộ GD&ĐT và các cơ quan đối tác, Bộ Y tế '
  'và Bộ LĐ-TB&XH, và UNICEF xem xét khi họ làm việc để thúc đẩy sức khoẻ tâm '
  'thần và hạnh phúc học sinh vị thành niên.', italic=True)

H('A. KHUYẾN NGHỊ CHO NGÀNH GIÁO DỤC (EDUCATION SECTOR RECOMMENDATIONS)',
  level=3)

P('1. Các cách tiếp cận toàn diện và toàn trường để cải thiện hạnh phúc học '
  'sinh và xây dựng khả năng phục hồi của học sinh.', bold=True)
P('Khí hậu trường học là bối cảnh cho việc học tập và phát triển của trẻ em và '
  'vị thành niên ở trường. Bối cảnh này bao gồm các yếu tố cơ bản như các mối '
  'quan hệ của học sinh với giáo viên và bạn bè, môi trường trường học, sự '
  'tham gia của học sinh, phụ huynh và giáo viên, và các trải nghiệm an toàn '
  'của học sinh ở trường. Một khí hậu trường học tích cực, hỗ trợ là thiết '
  'yếu cho việc học tập và hạnh phúc của học sinh. Các khuyến nghị cụ thể về '
  'khí hậu trường học bao gồm:')

P('b. Thúc đẩy các mối quan hệ giáo viên – học sinh tích cực. Các chương trình '
  'thúc đẩy sức khoẻ tâm thần dựa trên trường học phổ quát thấy rằng chất '
  'lượng các mối quan hệ giáo viên – học sinh dự báo các kết quả. Các bên '
  'liên quan đã chia sẻ các mô hình cải thiện các mối quan hệ giáo viên – '
  'học sinh, bao gồm tái cấu trúc lịch lớp học để học sinh gặp các giáo viên '
  'nhất định hai lần/ngày; tích hợp thời gian để học sinh – giáo viên chia '
  'sẻ và hỗ trợ vào chương trình giảng dạy chuẩn; ưu tiên các mối quan hệ '
  'giáo viên – học sinh sao cho giáo viên nhận được sự công nhận hoặc được '
  'đánh giá dựa trên chất lượng các mối quan hệ; và giảm sĩ số lớp.', bold=True)

P('c. Cấm sử dụng kỷ luật thể chất trong các trường học. Liên quan đến vấn '
  'đề hạnh phúc học sinh vị thành niên là nhu cầu thiết yếu loại bỏ kỷ luật '
  'thể chất trong các trường học. Các chính sách và thực tiễn phải giao tiếp '
  'rõ ràng rằng kỷ luật thể chất không được chấp nhận trong các lớp học. '
  'Quan trọng là giáo viên phải được trang bị các kỹ năng để tạo động lực '
  'cho học sinh và thực hiện kỷ luật tích cực, phi bạo lực.', bold=True)

P('d. Thúc đẩy sự tham gia và kết nối của học sinh với trường học. Hạnh phúc '
  'học sinh được cải thiện bởi sự gắn kết với bạn bè trong các hoạt động '
  'đồng chương trình và ngoại khoá, bao gồm nghệ thuật, thể thao, câu lạc bộ, '
  'v.v. Các hoạt động cung cấp cho học sinh cơ hội quý giá để phát triển các '
  'kỹ năng giải quyết vấn đề, giao tiếp, xây dựng mối quan hệ và khả năng '
  'phục hồi cảm xúc.', bold=True)

P('e. Thúc đẩy lòng tốt và các mối quan hệ bạn bè tích cực. Các trải nghiệm '
  'bắt nạt của học sinh đại diện cho một trong những thách thức lớn nhất '
  'của các em và có bằng chứng mạnh mẽ rằng bắt nạt đặt học sinh vào nguy '
  'cơ các vấn đề sức khoẻ tâm thần và các kết quả cuộc sống tiêu cực.',
  bold=True)

# --- Trang 79 ---
PageMark('--- Trang 79, UNICEF Việt Nam, 2022 ---')

P('f. Giảm áp lực học tập. Tất cả các bên liên quan đồng ý rằng học sinh đang '
  'chịu áp lực học tập đáng kể và các trải nghiệm áp lực học tập của học '
  'sinh, bao gồm áp lực học tập, lo lắng về điểm, chán nản liên quan đến '
  'học tập, tự kỳ vọng và khối lượng công việc, đều có liên quan đáng kể '
  'với các vấn đề sức khoẻ tâm thần học sinh. Các em gái đặc biệt dễ bị '
  'tổn thương với áp lực học tập. Rõ ràng rằng giảm áp lực học tập sẽ là '
  'một thách thức đáng kể vì nhiều yếu tố hệ thống, nền tảng đóng góp vào '
  'vấn đề này. Đào tạo và đánh giá giáo viên, lịch trình và trọng tâm kỳ '
  'thi, niềm tin về các mục tiêu và kết quả học tập, v.v. đều đóng góp vào '
  'văn hoá áp lực cao hiện tại. Bộ GD&ĐT có thể xem xét một con đường tiến '
  'tới, như một lực lượng đặc nhiệm được chỉ định để giải quyết vấn đề này, '
  'và bao gồm một quy trình đánh giá áp lực học tập trong mỗi phòng ban và '
  'hoạt động.', bold=True)

P('2. Các cách tiếp cận phổ quát trực tiếp nhắm vào sức khoẻ tâm thần và hạnh '
  'phúc.', bold=True)

P('a. Cải thiện hiểu biết sức khoẻ tâm thần của giáo viên, phụ huynh và học '
  'sinh. Các bên liên quan không có sự hiểu biết chung về sức khoẻ tâm thần '
  'học sinh. Toàn bộ cộng đồng trường học phải có kiến thức hơn về các vấn '
  'đề sức khoẻ tâm thần phổ biến và các triệu chứng, tỷ lệ lưu hành các vấn '
  'đề này ở học sinh, các yếu tố đặt học sinh vào nguy cơ các vấn đề và các '
  'chiến lược để giảm nguy cơ.', bold=True)

P('b. Dạy học sinh các kỹ năng cần thiết cho sức khoẻ tâm thần và hạnh phúc '
  'tích cực. Việc có được các kỹ năng xã hội và cảm xúc có liên quan đến '
  'phát triển thanh thiếu niên tích cực, thành tích học tập, các hành vi '
  'lối sống lành mạnh, và giảm trầm cảm và lo âu, bạo lực, bắt nạt, xung '
  'đột và tức giận. Các kỹ năng sống thiết yếu cho vị thành niên bao gồm '
  'điều tiết cảm xúc, giải quyết vấn đề, giao tiếp, giải quyết xung đột '
  'và duy trì các mối quan hệ lành mạnh.', bold=True)

P('i. Giáo dục giới tính có thể được cung cấp trong quan hệ đối tác với giáo '
  'dục về sức khoẻ tâm thần và hạnh phúc. Nhiều học sinh, giáo viên và phụ '
  'huynh bày tỏ lo ngại về các mối quan hệ tình dục. Các kỹ năng cho các '
  'mối quan hệ lãng mạn lành mạnh và tôn trọng (phi bạo lực) và sức khoẻ '
  'tình dục là cần thiết để thúc đẩy sức khoẻ và khả năng phục hồi của vị '
  'thành niên. Các chương trình này cũng nên bao gồm thông tin toàn diện, '
  'không phán xét về xu hướng tính dục và bản dạng giới để giảm phân biệt '
  'đối xử với học sinh LGBTQ.')

P('3. Cách tiếp cận chăm sóc theo bước nhắm đích để nhận diện và can thiệp '
  'sớm cho học sinh có các vấn đề sức khoẻ tâm thần.', bold=True)
P('Tất cả các bên liên quan đồng ý về nhu cầu cải thiện năng lực để nhận diện '
  'trẻ em có các vấn đề sớm và cung cấp cho các em hỗ trợ mà các em cần để '
  'giảm các triệu chứng sức khoẻ tâm thần và lấy lại sức khoẻ và chức năng '
  'đầy đủ.')

# --- Trang 80 ---
PageMark('--- Trang 80, UNICEF Việt Nam, 2022 ---')

P('a. Cung cấp cho việc nhận diện sớm các vấn đề sức khoẻ tâm thần học sinh. '
  'Điều này bắt đầu với cải thiện hiểu biết sức khoẻ tâm thần giữa giáo '
  'viên và phụ huynh, tuy nhiên kiến thức đơn lẻ là không đủ. Một chương '
  'trình sàng lọc được khuyến nghị để nhận diện nhất quán hơn các học sinh '
  'có nguy cơ các vấn đề.', bold=True)

P('i. Chương trình Sức khoẻ Học đường giai đoạn 2021–2025 vừa được thiết lập '
  '(Quyết định Thủ tướng số 1660/QĐ-TTg, ngày 2/10/2021 và Quyết định Thủ '
  'tướng số 85/QĐ-TTg, ngày 17/1/2022) có thể cung cấp hỗ trợ chính sách '
  'để mở rộng vai trò của nhân viên y tế học đường bao gồm sàng lọc phát '
  'hiện sớm các vấn đề sức khoẻ tâm thần ở học sinh vị thành niên.')

P('ii. Với hướng dẫn của Bộ GD&ĐT, các Sở GD&ĐT sẽ cần xác định các nguồn lực '
  'địa phương cho điều trị sức khoẻ tâm thần chuyên nghiệp (ví dụ: bệnh viện '
  'hoặc khoa tâm thần, phòng khám tâm lý, trung tâm bảo trợ xã hội cung cấp '
  'dịch vụ tư vấn vị thành niên). Các con đường giới thiệu cho mỗi huyện nên '
  'được phát triển và chia sẻ với tất cả các trường trung học cơ sở và phổ '
  'thông trong huyện.')

P('b. Cung cấp tư vấn học đường chuyên nghiệp cho học sinh có các vấn đề sức '
  'khoẻ tâm thần mức trung bình. Một vị trí chính thức cho Tư vấn viên Học '
  'đường phải được thiết lập. Các tư vấn viên học đường nên có đào tạo về '
  'phát triển vị thành niên, sức khoẻ tâm thần vị thành niên, đánh giá và '
  'lập kế hoạch điều trị sức khoẻ tâm thần cơ bản, các can thiệp nhóm và cá '
  'nhân dựa trên bằng chứng cho các vấn đề sức khoẻ tâm thần phổ biến, điều '
  'phối với giáo viên và phụ huynh trong hỗ trợ học sinh, và giới thiệu cho '
  'học sinh có các vấn đề nghiêm trọng cho điều trị y tế và tâm lý chuyên '
  'gia.', bold=True)

P('4. Xem xét các nhu cầu cụ thể của học sinh dân tộc thiểu số, các em gái '
  'và học sinh LGBTQ trong phát triển chính sách và chương trình.', bold=True)
P('Các nhóm này có nguy cơ cao hơn về các vấn đề sức khoẻ tâm thần. Các chương '
  'trình có thể điều chỉnh cách tiếp cận tiếp cận, thu hút và thực hiện đối '
  'với học sinh thuộc các dân tộc, giới và bản dạng giới/tính dục khác nhau '
  'để tối đa hoá thành công chương trình với các nhóm dễ tổn thương này.')

P('5. Xây dựng nguồn nhân lực.', bold=True)
P('Nguồn nhân lực cho sức khoẻ tâm thần học sinh vị thành niên đang thiếu '
  'trong tất cả các ngành. Bộ GD&ĐT có thể xem xét các nhu cầu nguồn nhân '
  'lực và phát triển các chính sách để xây dựng năng lực trong ngành giáo '
  'dục.')

# --- Trang 81 ---
doc.add_page_break()
PageMark('--- Trang 81, UNICEF Việt Nam, 2022 ---')

H('B. KHUYẾN NGHỊ CHO NGÀNH Y TẾ (HEALTH SECTOR RECOMMENDATIONS)', level=3)

P('1. Xây dựng nguồn nhân lực. Hệ thống y tế có rất ít bác sĩ và y tá chuyên '
  'về sức khoẻ tâm thần trẻ em và vị thành niên.', bold=True)

P('a. Cải thiện kiến thức và kỹ năng cho tất cả nhân viên chăm sóc sức khoẻ '
  'bằng cách bao gồm nội dung khoá học liên quan đến sức khoẻ tâm thần trẻ '
  'em và vị thành niên trong chương trình giảng dạy trường y đa khoa.')

P('b. Phát triển nguồn nhân lực sức khoẻ tâm thần trẻ em và vị thành niên '
  'chuyên biệt. Các thách thức đáng kể đóng góp vào sự thiếu hụt các '
  'chuyên gia tâm thần tại Việt Nam. Lãnh đạo Bộ Y tế nên trực tiếp giải '
  'quyết các rào cản này với các chính sách và chương trình rõ ràng được '
  'thiết kế để cải thiện năng lực khu vực y tế chăm sóc cho tất cả những '
  'người mắc bệnh tâm thần, bao gồm trẻ em và vị thành niên.')

P('2. Đào tạo nhân viên y tế học đường trong sức khoẻ tâm thần vị thành '
  'niên cơ bản. Phạm vi đào tạo này sẽ phụ thuộc vào các mô hình hợp tác '
  'liên ngành được các Bộ phát triển để giải quyết nhu cầu sức khoẻ tâm '
  'thần học sinh.', bold=True)

P('3. Xây dựng hợp tác Sở Y tế – Sở GD&ĐT. Hợp tác chính thức giữa Sở Y '
  'tế và Sở GD&ĐT là cần thiết để giải quyết sức khoẻ tâm thần học sinh '
  'vị thành niên.', bold=True)

H('C. KHUYẾN NGHỊ CHO NGÀNH XÃ HỘI (SOCIAL SECTOR RECOMMENDATIONS)', level=3)

P('1. Xây dựng nguồn nhân lực.', bold=True)
P('a. Cải thiện kiến thức và kỹ năng cho tất cả nhân viên xã hội bằng cách '
  'bao gồm nội dung khoá học liên quan đến sức khoẻ tâm thần trẻ em và vị '
  'thành niên trong chương trình giảng dạy công tác xã hội chung.')
P('b. Phát triển nguồn nhân lực sức khoẻ tâm thần vị thành niên và gia đình '
  'chuyên biệt.')

P('2. Xây dựng hợp tác Sở LĐ-TB&XH – Sở GD&ĐT. Hợp tác chính thức giữa Sở '
  'LĐ-TB&XH và Sở GD&ĐT là cần thiết để giải quyết sức khoẻ tâm thần học '
  'sinh vị thành niên. Các mô hình hợp tác tiềm năng bao gồm:', bold=True)

# --- Trang 82 ---
PageMark('--- Trang 82, UNICEF Việt Nam, 2022 ---')

P('a. Thiết lập quan hệ đối tác chính thức giữa Sở LĐ-TB&XH và Sở GD&ĐT '
  'để giao tiếp và chia sẻ thường xuyên dữ liệu về tỷ lệ các vấn đề sức '
  'khoẻ tâm thần vị thành niên và các can thiệp tiềm năng.')
P('b. Tích hợp nội dung đào tạo sức khoẻ tâm thần vị thành niên vào các '
  'chương trình đào tạo và hỗ trợ phụ huynh hiện có của MOLISA/DOLISA.')
P('c. Hợp tác Sở LĐ-TB&XH với Sở GD&ĐT để cung cấp các khoá kỹ năng sống '
  'dựa trên trường học cho học sinh.')
P('d. Đồng phát triển các con đường chăm sóc cho học sinh có các vấn đề '
  'sức khoẻ tâm thần được giới thiệu đến các Trung tâm Công tác Xã hội '
  'để tư vấn cho trẻ em và gia đình.')

H('D. KHUYẾN NGHỊ CHO UNICEF (UNICEF RECOMMENDATIONS)', level=3)

P('UNICEF đã hỗ trợ nghiên cứu và phát triển về sức khoẻ tâm thần trẻ em '
  'tại Việt Nam trong nhiều năm và ở vị trí xuất sắc để hỗ trợ các phát '
  'triển chính sách và chương trình liên quan đến nghiên cứu này trong '
  'tương lai. Các khuyến nghị cụ thể cho UNICEF bao gồm:', italic=True)

P('1. Vận động công chúng và nâng cao nhận thức.', bold=True)
P('UNICEF đang ở vị trí then chốt để sử dụng các phát hiện nghiên cứu nhằm '
  'nâng cao nhận thức về tầm quan trọng của sức khoẻ tâm thần vị thành '
  'niên tại Việt Nam.')

P('2. Tạo thuận lợi cho hợp tác liên ngành về phát triển chính sách và '
  'chương trình học sinh vị thành niên.', bold=True)
P('Các nhà quản lý tham gia và các chuyên gia bộ ngành từ cả ba ngành đã '
  'xác định sự thiếu điều phối liên ngành là một rào cản đáng kể đối với '
  'tiến bộ, và xác định nhu cầu cụ thể cho một hội nghị hoặc cuộc họp '
  'liên ngành về chủ đề này.')

P('3. Vận động và hỗ trợ nghiên cứu thêm.', bold=True)
P('Các lĩnh vực cụ thể được xác định bao gồm các yếu tố nguy cơ cho các '
  'nhóm học sinh vị thành niên dễ bị tổn thương, bao gồm các em gái, học '
  'sinh dân tộc thiểu số và học sinh LGBTQ.')

P('4. Hỗ trợ sự phù hợp của các vấn đề bảo vệ trẻ em và sức khoẻ tâm thần '
  'vị thành niên.', bold=True)

# --- Trang 83 ---
PageMark('--- Trang 83, UNICEF Việt Nam, 2022 ---')
P('UNICEF có chuyên môn và kinh nghiệm lớn về bảo vệ trẻ em tại Việt Nam. '
  'Các vấn đề bảo vệ trẻ em và sức khoẻ tâm thần vị thành niên chia sẻ '
  'một số yếu tố nguy cơ và con đường chung.')

P('5. Tích hợp các vấn đề sức khoẻ tâm thần vị thành niên vào các chương '
  'trình UNICEF liên quan.', bold=True)

H('E. KHUYẾN NGHỊ NGHIÊN CỨU THÊM (RECOMMENDATIONS FOR FURTHER RESEARCH)',
  level=3)

P('Có đủ bằng chứng về nhu cầu sức khoẻ tâm thần học sinh vị thành niên '
  'Việt Nam, các hạn chế của các hệ thống hỗ trợ hiện tại, và các chương '
  'trình có triển vọng để cải thiện sức khoẻ tâm thần và hạnh phúc của '
  'học sinh. Lập kế hoạch chính sách và chương trình nên tiến triển. Tuy '
  'nhiên, các mô hình hợp tác và chương trình nên được thí điểm trong '
  'các trường, huyện hoặc tỉnh trước khi mở rộng quy mô đến cấp quốc '
  'gia. Nghiên cứu nên đánh giá các chương trình thí điểm này và hướng '
  'dẫn các nỗ lực mở rộng tiếp theo.')

P('Nghiên cứu về nhu cầu sức khoẻ tâm thần cụ thể của học sinh vị thành '
  'niên từ các nhóm dễ bị tổn thương đang thiếu. Các nghiên cứu về học '
  'sinh dân tộc thiểu số và LGBTQ là cần thiết để hiểu rõ hơn các vấn '
  'đề các học sinh này trải qua, các rào cản các em đối mặt, và cách các '
  'trường và cộng đồng có thể hỗ trợ tốt nhất sức khoẻ và hạnh phúc của '
  'các em.')

H('F. HỖ TRỢ ĐẠI DỊCH (PANDEMIC SUPPORT)', level=3)

P('Học sinh trên khắp thế giới đã phải chịu đựng trong đại dịch COVID-19. '
  'Đóng cửa trường học và các hạn chế xã hội đã tước đoạt học sinh các '
  'mối quan hệ xã hội thiết yếu cho phát triển lành mạnh. Stress tài '
  'chính trên các gia đình đã tác động đến sức khoẻ tâm thần học sinh. '
  'Sự mất gắn kết và stress đã dẫn đến tỷ lệ trầm cảm và lo âu ngày càng '
  'tăng ở người trẻ trên toàn cầu. Việt Nam không phải là ngoại lệ và '
  'tất cả các bên liên quan bày tỏ lo ngại về tác động của đại dịch lên '
  'sức khoẻ và phát triển vị thành niên.')

P('Trong khi các khuyến nghị này chủ yếu tập trung vào các phát triển '
  'chính sách và chương trình dài hạn, rõ ràng là học sinh vị thành niên '
  'cần hỗ trợ sức khoẻ tâm thần và hạnh phúc ngay bây giờ. Bộ GD&ĐT và '
  'tất cả các bên liên quan phải xem xét các can thiệp kịp thời để hỗ '
  'trợ học sinh bị tác động bởi đại dịch.', bold=True)

doc.add_page_break()

# ============================================================
# PHU LUC 1 — HUONG DAN FGD
# ============================================================
H('PHỤ LỤC 1: HƯỚNG DẪN VÀ CÂU HỎI THẢO LUẬN NHÓM TẬP TRUNG (FGD)', level=1)
P('(Appendix 1 — Focus Group Guidelines and Questions)', italic=True, align='center', size=11)

# --- Trang 84 ---
PageMark('--- Trang 84, UNICEF Việt Nam, 2022 ---')
H('Hướng dẫn FGD với Học sinh (Guidelines for Focus Groups with Students)', level=2)

P('** Chúng tôi thực sự đánh giá cao sự tham gia của các em vào nhóm này. Hôm nay '
  'chúng tôi chỉ muốn có một cuộc thảo luận về cuộc sống, trường học và một số điều '
  'có thể gây stress cho các em. Chúng tôi đang nói chuyện với các em như các em từ '
  '5 tỉnh khác nhau tại Việt Nam. Theo cách này, chúng tôi có thể tìm hiểu thêm về '
  'cuộc sống của vị thành niên như các em, và có thể hỗ trợ học sinh tốt hơn ở '
  'trường. Cuộc thảo luận này đang được ghi âm chỉ để chúng tôi không bỏ lỡ bất cứ '
  'điều gì các em nói. Chúng tôi sẽ giữ tất cả mọi thứ được nói ẩn danh — điều đó '
  'có nghĩa là chúng tôi sẽ không bao giờ sử dụng tên hoặc bản dạng của các em khi '
  'chia sẻ những điều chúng tôi học được từ cuộc thảo luận này.', italic=True)
P('** Mẫu đồng ý tham gia (assent form).', italic=True)

P('1. Các em thích trường học như thế nào? Các em thích gì và không thích gì?')
P('2. Điều gì gây stress cho các em trong cuộc sống? (nếu không được đề cập, hỏi '
  'cụ thể về stress học tập, stress về các mối quan hệ bạn bè, stress về các vấn '
  'đề hoặc hoàn cảnh gia đình, internet/mạng xã hội/bắt nạt mạng)')
P('3. Stress ảnh hưởng đến cuộc sống của các em như thế nào? (Có giúp các em không? '
  'Có gây vấn đề cho các em không? Có giúp các em với bài tập ở trường hay khiến '
  'việc học khó khăn không?)')
P('4. Các em làm gì sau giờ học? Các em làm gì vào cuối tuần? (tìm hiểu xem các em '
  'có thời gian nào để thư giãn hoặc tham gia các môn thể thao, sở thích hoặc các '
  'hoạt động vui chơi không)')
P('5. Các em làm gì để giúp bản thân khi bị stress?')
P('6. Nếu học sinh đang gặp khó khăn về cảm xúc, do stress hoặc các vấn đề khác, '
  'giáo viên hoặc nhà trường có giúp đỡ theo cách nào không?')
P('7. Các em nghĩ trường có thể làm gì để hỗ trợ tốt hơn cho học sinh về mặt cảm xúc?')

# --- Trang 85 ---
PageMark('--- Trang 85, UNICEF Việt Nam, 2022 ---')
H('Hướng dẫn FGD với Học sinh LGBTQ (Guidelines for Focus Groups with LGBTQ Students)',
  level=2)
P('(Tương tự hướng dẫn chung với học sinh, có thêm các câu hỏi đặc thù về '
  'việc là LGBTQ ảnh hưởng đến cuộc sống vị thành niên, trải nghiệm stress, '
  'phân biệt đối xử, bắt nạt và lo ngại về tương lai.)')

P('1. Các em thích trường học như thế nào?')
P('2. Việc là LGBTQ ảnh hưởng đến cuộc sống tuổi teen của các em như thế nào? (Bản '
  'dạng tích cực? Sợ bị kỳ thị hoặc phân biệt đối xử? Lo lắng về phản ứng của gia '
  'đình?)')
P('3. Việc là LGBTQ ảnh hưởng đến cuộc sống học sinh của các em như thế nào? (Lo '
  'lắng về phản ứng bạn bè? Trải qua phân biệt đối xử và/hoặc bắt nạt từ bạn bè? '
  'Lo lắng về phản ứng giáo viên?)')
P('4. Các em cảm thấy thế nào khi nghĩ về tương lai? (Tích cực? Lo lắng hoặc bi '
  'quan?)')
P('5. Stress trong cuộc sống ảnh hưởng đến các em như thế nào?')
P('6. Các em làm gì để giúp bản thân khi bị stress?')
P('7. Nếu các em hoặc bạn bè đang gặp khó khăn về cảm xúc, giáo viên hoặc trường '
  'có giúp đỡ theo cách nào không?')
P('8. Các em nghĩ trường có thể làm gì để hỗ trợ tốt hơn các em và các học sinh '
  'LGBTQ khác về mặt cảm xúc và học tập?')

# --- Trang 86 ---
PageMark('--- Trang 86, UNICEF Việt Nam, 2022 ---')
H('Hướng dẫn FGD với Phụ huynh (Guidelines for Focus Groups with Parents)', level=2)
P('1. Con của anh/chị đang làm thế nào về mặt cảm xúc? Hạnh phúc của cháu như thế nào?')
P('2. Con của anh/chị có trải qua stress không? Bao nhiêu stress/bao thường xuyên?')
P('3. Điều gì gây stress cho con anh/chị trong cuộc sống?')
P('4. Stress ảnh hưởng đến cuộc sống của con anh/chị như thế nào?')
P('5. Con anh/chị làm gì sau giờ học? Làm gì vào cuối tuần?')
P('6. Anh/chị có nghĩ trẻ em dành quá nhiều thời gian cho bài vở và học tập không?')
P('7. Trường và xã hội nên làm gì để giảm stress học sinh và hỗ trợ học sinh tốt hơn?')

H('Hướng dẫn FGD với Giáo viên (Guidelines for Focus Groups with Teachers)', level=2)
# --- Trang 87 ---
PageMark('--- Trang 87, UNICEF Việt Nam, 2022 ---')
P('1. Học sinh của anh/chị đang làm thế nào về mặt cảm xúc? Hạnh phúc của các em '
  'như thế nào?')
P('2. Học sinh của anh/chị có trải qua stress không? Bao nhiêu stress/bao thường '
  'xuyên?')
P('3. Điều gì gây stress cho học sinh trong cuộc sống?')
P('4. Stress ảnh hưởng đến học sinh trong cuộc sống như thế nào?')
P('5. Stress có giúp học sinh với bài vở trường học hay khiến các em khó thành '
  'công học tập hơn không?')
P('6. Học sinh thường làm bao nhiêu bài vở trường học ngoài giờ học?')
P('7. Bao nhiêu học sinh của anh/chị tham gia các lớp học thêm hoặc có gia sư sau '
  'giờ học hoặc vào cuối tuần?')
P('8. Anh/chị nghĩ gì về kỳ vọng của phụ huynh đối với thành tích học tập của '
  'con họ?')
P('9. Anh/chị nghĩ gì về kỳ vọng của học sinh đối với bản thân?')
P('10. Anh/chị có nghĩ trẻ em dành quá nhiều thời gian cho bài vở và học tập '
  'không?')
P('11. Trường có chương trình hoặc dịch vụ nào để hỗ trợ học sinh đang bị stress '
  'hoặc các vấn đề sức khoẻ tâm thần không?')
P('12. Anh/chị nghĩ trường và xã hội còn có thể làm gì khác để giảm stress học '
  'sinh và hỗ trợ học sinh tốt hơn?')

doc.add_page_break()

# ============================================================
# PHU LUC 2 — HUONG DAN KII
# ============================================================
H('PHỤ LỤC 2: HƯỚNG DẪN VÀ CÂU HỎI PHỎNG VẤN NGƯỜI CUNG CẤP THÔNG TIN CHỦ CHỐT (KII)',
  level=1)
P('(Appendix 2 — Key Informant Interview Guidelines and Questions)',
  italic=True, align='center', size=11)

# --- Trang 88-89 ---
PageMark('--- Trang 88, UNICEF Việt Nam, 2022 ---')
H('Hướng dẫn KII Bộ Y tế (Ministry of Health KII Guidelines)', level=2)

P('Người tham gia: Người cung cấp thông tin chủ chốt của Bộ Y tế và Sở Y tế '
  '(45–60 phút mỗi cuộc).', italic=True)
P('Mục tiêu: Khám phá nhận thức về nhu cầu sức khoẻ tâm thần học sinh và các yếu '
  'tố nguy cơ, hiểu biết về các chính sách liên quan, và kiến thức và hiểu biết '
  'về các chương trình MHPSS hiện có.', italic=True)

P('Thông tin người cung cấp thông tin chủ chốt:', bold=True)
P('Tên: ___________ | Ngày: ___________ | Chức danh: ___________')
P('Số năm kinh nghiệm về Giáo dục, Y tế hoặc Hỗ trợ Xã hội Vị thành niên: ________')

P('Các câu hỏi phỏng vấn bán cấu trúc:', bold=True)
P('1. Kinh nghiệm của anh/chị về chủ đề sức khoẻ tâm thần và hạnh phúc học sinh '
  'vị thành niên là gì?')
P('2. Vấn đề sức khoẻ tâm thần và hạnh phúc học sinh vị thành niên quan trọng '
  'như thế nào?')
P('3. Các vấn đề hoặc nhu cầu sức khoẻ tâm thần phổ biến ở học sinh vị thành '
  'niên Việt Nam là gì?')
P('4. Các yếu tố chính đặt trẻ em vào nguy cơ các vấn đề sức khoẻ tâm thần là '
  'gì? (xem xét áp lực học tập, gánh nặng học tập, khí hậu trường học, hoàn '
  'cảnh gia đình, lạm dụng chất, đói nghèo, bệnh tật, các vấn đề hệ thống '
  'trong giáo dục, y tế, hỗ trợ xã hội)')
P('5. Có cách nào mà ngành y tế (Bộ Y tế, Sở Y tế) hỗ trợ sức khoẻ tâm thần và '
  'hạnh phúc của trẻ em không?')
P('   a. Chính sách Bộ Y tế/Sở Y tế nào tồn tại về sức khoẻ tâm thần và hạnh '
  'phúc vị thành niên?')
P('   b. Dịch vụ hoặc can thiệp cụ thể nào tồn tại cho sức khoẻ cảm xúc và tâm '
  'thần của học sinh?')
P('   c. Các dịch vụ/can thiệp này có hiệu quả không?')
# --- Trang 89 ---
PageMark('--- Trang 89, UNICEF Việt Nam, 2022 ---')
P('6. Có sự điều phối nào giữa ngành y tế và ngành giáo dục không? Vui lòng mô '
  'tả.')
P('7. Các loại điều phối nào giữa hệ thống y tế và hệ thống trường học sẽ cải '
  'thiện hỗ trợ cho sức khoẻ tâm thần vị thành niên?')

H('Hướng dẫn KII Bộ LĐ-TB&XH (MOLISA KII Guidelines)', level=2)

P('Người tham gia: Người cung cấp thông tin chủ chốt của Bộ LĐ-TB&XH và Sở LĐ-TB&XH '
  '(45–60 phút mỗi cuộc).', italic=True)
P('Mục tiêu: Khám phá nhận thức về nhu cầu sức khoẻ tâm thần học sinh và các yếu '
  'tố nguy cơ, hiểu biết về các chính sách liên quan, và kiến thức và hiểu biết '
  'về các chương trình MHPSS hiện có.', italic=True)

P('Thông tin người cung cấp thông tin chủ chốt:', bold=True)
P('Tên: ___________ | Ngày: ___________ | Chức danh: ___________')
P('Số năm kinh nghiệm về Giáo dục, Y tế hoặc Hỗ trợ Xã hội Vị thành niên: ________')

P('Các câu hỏi phỏng vấn bán cấu trúc:', bold=True)
P('1. Kinh nghiệm của anh/chị về chủ đề sức khoẻ tâm thần và hạnh phúc học sinh '
  'vị thành niên là gì?')
P('2. Vấn đề sức khoẻ tâm thần và hạnh phúc học sinh vị thành niên quan trọng '
  'như thế nào?')
P('3. Các vấn đề hoặc nhu cầu sức khoẻ tâm thần phổ biến ở học sinh vị thành '
  'niên Việt Nam là gì?')
P('4. Các yếu tố chính đặt trẻ em vào nguy cơ các vấn đề sức khoẻ tâm thần là '
  'gì? (xem xét áp lực học tập, gánh nặng học tập, khí hậu trường học, hoàn '
  'cảnh gia đình, lạm dụng chất, đói nghèo, các vấn đề hệ thống trong giáo '
  'dục, y tế, hỗ trợ xã hội)')
P('5. Có cách nào mà khu vực dịch vụ xã hội (Bộ LĐ-TB&XH, Sở LĐ-TB&XH) hỗ trợ '
  'sức khoẻ tâm thần và hạnh phúc của trẻ em không?')
P('   a. Chính sách Bộ LĐ-TB&XH/Sở LĐ-TB&XH nào tồn tại về sức khoẻ tâm thần '
  'và hạnh phúc vị thành niên?')
P('   b. Dịch vụ hoặc can thiệp cụ thể nào tồn tại cho sức khoẻ cảm xúc và '
  'tâm thần của học sinh?')
P('   c. Các dịch vụ/can thiệp này có hiệu quả không?')
P('6. Có sự điều phối nào với ngành giáo dục không? Vui lòng mô tả.')
P('7. Các loại điều phối nào với hệ thống trường học sẽ cải thiện hỗ trợ cho '
  'sức khoẻ tâm thần vị thành niên?')

# --- Trang 90 ---
PageMark('--- Trang 90, UNICEF Việt Nam, 2022 ---')
H('Hướng dẫn KII Ngành Giáo dục (Education Sector KII Guidelines)', level=2)
P('Người tham gia: Hiệu trưởng, Sở GD&ĐT, và Bộ GD&ĐT (45–60 phút mỗi cuộc).',
  italic=True)

P('1. Kinh nghiệm của anh/chị về chủ đề sức khoẻ tâm thần và hạnh phúc học '
  'sinh vị thành niên là gì?')
P('2. Vấn đề sức khoẻ tâm thần và hạnh phúc học sinh vị thành niên quan trọng '
  'như thế nào?')
P('3. Các vấn đề hoặc nhu cầu sức khoẻ tâm thần phổ biến ở học sinh vị thành '
  'niên Việt Nam là gì?')
P('4. Có cách nào mà các trường và/hoặc hệ thống giáo dục đặt trẻ em vào nguy '
  'cơ các vấn đề sức khoẻ tâm thần không? Những khía cạnh cụ thể nào của '
  'trường học đóng góp vào các vấn đề xã hội, cảm xúc và tâm lý của học sinh?')
P('5. Có cách nào mà các trường và/hoặc hệ thống giáo dục hỗ trợ sức khoẻ tâm '
  'thần và hạnh phúc của trẻ em không?')
P('   a. Chính sách trường nào tồn tại về sức khoẻ tâm thần và hạnh phúc '
  'học sinh?')
P('   b. Dịch vụ hoặc can thiệp cụ thể nào các trường có cho sức khoẻ cảm xúc '
  'và tâm thần của học sinh?')
P('   c. Các dịch vụ/can thiệp này có hiệu quả không?')
P('6. Các trường còn có thể làm gì khác để hỗ trợ sức khoẻ tâm thần và hạnh '
  'phúc của trẻ em?')

doc.add_page_break()

# ============================================================
# PHU LUC 3 — DU LIEU DINH LUONG
# ============================================================
H('PHỤ LỤC 3: KẾT QUẢ DỮ LIỆU ĐỊNH LƯỢNG', level=1)
P('(Appendix 3 — Quantitative Data Results)', italic=True, align='center', size=11)

# --- Trang 92 ---
PageMark('--- Trang 92, UNICEF Việt Nam, 2022 ---')

H('Dữ liệu mô tả: Người tham gia học sinh (n = 668)', level=2)

P('Tuổi: Trung bình 14,23 | Median = 14 | SD = 1,9 | Dải: 11–18 tuổi')
P('Giới tính: Nam 228 (34 %) | Nữ 439 (66 %)')

P('Dân tộc:', bold=True)
P('Kinh: 307 (45,9 %) | Gia-rai: 110 (16,5 %) | Thái: 98 (14,7 %) | Mông: 22 '
  '(3,3 %) | Ba Na: 13 (1,9 %) | Mường: 5 (0,7 %) | Khơ Mú: 5 (0,7 %) | Khác: '
  '24 (Cống, Dao, Ê đê, Hàn Quốc, Lào, Nhắng, Nùng, Phù Lá, Sơ rá, Tày, Xơ '
  'đăng) | Không xác định: 85 (12,7 %)')

P('Tỉnh: Hà Nội 124 (18,6 %) | Điện Biên 184 (27,5 %) | Đồng Tháp 178 (26,6 %) '
  '| Gia Lai 176 (26,3 %)')

P('Loại trường: Công lập 608 | Bán công 60', italic=True)
P('Cấp học: THCS 425 (63,6 %) | THPT 242 (36,2 %)')
P('Khối lớp: 6–12 | Khối 7 chiếm 25,7 % | Khối 6: 18,6 % | Khối 8: 19,3 % | '
  'Khối 9: 15 % | Khối 10: 16,5 % | Khối 11: 4,9 %')

# --- Trang 93-94: SDQ-25 ---
PageMark('--- Trang 93–94, UNICEF Việt Nam, 2022 ---')

H('Triệu chứng sức khoẻ tâm thần (SDQ-25)', level=2)

P('1. Sức khoẻ tâm thần tổng thể học sinh. 26,1 % học sinh vị thành niên báo '
  'cáo các triệu chứng liên quan đến nguy cơ trung bình hoặc cao về các vấn '
  'đề sức khoẻ tâm thần. Các loại vấn đề phổ biến nhất được báo cáo là các '
  'vấn đề bạn bè (32 %) và các vấn đề cảm xúc (30,9 %). Các triệu chứng tăng '
  'động được báo cáo bởi 14,4 % học sinh và các vấn đề hành vi được tìm thấy '
  'ở 11 % học sinh.', bold=True)

P('Bảng A3.1 — Mức nguy cơ học sinh (N = 668):', bold=True, italic=True)
t = doc.add_table(rows=6, cols=4)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
set_grid(t, [5.0, 3.5, 3.5, 3.5])
rows_sdq = [
    ('', 'Bình thường N (%)', 'Nguy cơ Trung bình N (%)', 'Nguy cơ Cao N (%)'),
    ('Tổng vấn đề SKTT', '439 (65,7 %)', '112 (16,8 %)', '62 (9,3 %)'),
    ('Vấn đề hành vi', '581 (87 %)', '47 (7 %)', '27 (4 %)'),
    ('Vấn đề cảm xúc', '446 (66,8 %)', '104 (15,6 %)', '102 (15,3 %)'),
    ('Tăng động', '552 (82,6 %)', '66 (9,9 %)', '30 (4,5 %)'),
    ('Vấn đề bạn bè', '436 (65,3 %)', '167 (25,0 %)', '47 (7,0 %)'),
]
for ri, rd in enumerate(rows_sdq):
    for ci, v in enumerate(rd):
        c = t.rows[ri].cells[ci]; c.text = str(v)
        colw(c, [5.0, 3.5, 3.5, 3.5][ci])
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if ri == 0: r.bold = True
        if ri == 0: shade(c, 'D9E2F3')

P('2. Sức khoẻ tâm thần và giới tính: Tỷ lệ tương tự giữa các em trai và em '
  'gái ngoại trừ Vấn đề Cảm xúc (triệu chứng trầm cảm và lo âu), nơi chúng '
  'ta thấy tỷ lệ cao hơn đáng kể ở các em gái (M = 4,92) so với các em trai '
  '(M = 4,24), p < 0,01.')

P('3. Sức khoẻ tâm thần và cấp học: Khác biệt đáng kể với học sinh THPT có '
  'nhiều vấn đề hành vi, tăng động, cảm xúc và tổng vấn đề hơn so với THCS. '
  'Cụ thể tổng vấn đề: THCS M = 12,62 vs THPT M = 14,27 (p < 0,01).')

# --- Trang 95-99: School Climate + Academic Pressure ---
PageMark('--- Trang 95–99, UNICEF Việt Nam, 2022 ---')

H('Khí hậu trường học (MDS3) và sức khoẻ tâm thần', level=2)

P('Mối liên hệ yếu – vừa giữa báo cáo của học sinh về các yếu tố khí hậu '
  'trường học và sức khoẻ tâm thần học sinh. Nhận thức của học sinh về an '
  'toàn, gắn kết và môi trường trường học đều có tương quan đáng kể với tất '
  'cả các vấn đề sức khoẻ tâm thần. Có mối quan hệ mạnh vừa phải giữa Tổng '
  'Khó khăn SKTT và gắn kết trường học cảm nhận của học sinh (r = −0,22, '
  'p < 0,01) và môi trường trường học (r = −0,22, p < 0,01).')

P('Khí hậu trường học và giới tính: Các em trai và em gái báo cáo nhận thức '
  'tương tự về an toàn và môi trường. Các em trai đánh giá gắn kết cao hơn '
  'các em gái (M = 68,16 vs 66,24, p < 0,05).')

P('Khí hậu trường học và cấp học: THPT nhận thức an toàn hơn (M = 15,14 vs '
  '13,56, p < 0,01), nhưng THCS gắn kết cao hơn (M = 67,84 vs 65,33, p < '
  '0,01). Bắt nạt cao hơn ở THCS; sử dụng ma tuý và rượu cao hơn ở THPT.')

P('Khí hậu trường học theo tỉnh:', bold=True)
P('• An toàn: Hà Nội cao nhất (M = 17,86), Gia Lai thấp nhất (M = 11,19), '
  'khác biệt đáng kể.')
P('• Gắn kết: Đồng Tháp cao nhất (M = 68,98), Điện Biên thấp nhất (M = 65,75).')

H('Áp lực học tập (ESSA)', level=2)

P('Thang đo ESSA bao gồm 16 mục, 5 khía cạnh: áp lực học tập (4 mục), lo lắng '
  'về điểm (3 mục), chán nản (3 mục), tự kỳ vọng (3 mục), khối lượng công '
  'việc (3 mục).')

P('Áp lực học tập và sức khoẻ tâm thần: Mối quan hệ vừa – mạnh giữa các '
  'trải nghiệm áp lực học tập và SKTT. Tất cả 5 khía cạnh đều có tương '
  'quan đáng kể với tất cả các vấn đề SKTT (trừ khối lượng công việc × vấn '
  'đề bạn bè). Tổng áp lực học tập × Tổng khó khăn SKTT: r = 0,42 (p < 0,01).')

P('Áp lực học tập và giới tính: Các em gái trải qua mức áp lực học tập cao '
  'hơn đáng kể (M = 48,11 vs 44,27, p < 0,01), bao gồm áp lực học tập, lo '
  'lắng về điểm, tự kỳ vọng và chán nản. Khối lượng công việc tương tự nam '
  '(7,92) và nữ (8,03).')

P('Áp lực học tập và cấp học: THPT báo cáo áp lực học tập nhiều hơn đáng '
  'kể so với THCS (M = 49,82 vs 44,93, p < 0,01).')

P('Áp lực học tập theo tỉnh: Cao hơn ở Gia Lai (M = 48,71) và Hà Nội (M = '
  '47,61) so với Đồng Tháp (45,40) và Điện Biên (45,73).')

# --- Trang 100-101 ---
PageMark('--- Trang 100–101, UNICEF Việt Nam, 2022 ---')

H('Thời gian học và học thêm', level=2)

P('Thời gian học sau giờ học:', bold=True)
P('• Hầu như không có: 24 (3,6 %) | Dưới 1 giờ: 60 (9,0 %) | 1–2 giờ: 225 '
  '(33,7 %) | 2–3 giờ: 169 (25,3 %) | Hơn 3 giờ: 188 (28,1 %)')
P('→ Hơn 50 % học sinh học > 2 giờ/ngày, 28 % học hơn 3 giờ/đêm.')

P('Thời gian học thêm/dạy kèm:', bold=True)
P('• Hầu như không có: 213 (31,9 %) | 1–2 giờ: 138 (20,7 %) | 3–5 giờ: 130 '
  '(19,5 %) | 6–8 giờ: 83 (12,4 %) | 9–12 giờ: 50 (7,5 %) | > 12 giờ: 51 '
  '(7,6 %)')
P('→ Khoảng 50 % có ≤ 2 giờ học thêm/tuần, 50 % khác có > 3 giờ/tuần, bao '
  'gồm 15 % có > 9 giờ/tuần.')

H('Thành tích học tập (GPA tự báo cáo)', level=2)
P('Toán: Dưới 5: 53 (7,9 %) | 5–6,9: 199 (29,8 %) | 7–7,9: 152 (22,8 %) | '
  '8–10: 261 (39,1 %)')
P('Ngữ Văn: Dưới 5: 19 (2,8 %) | 5–6,9: 172 (25,7 %) | 7–7,9: 252 (37,7 %) | '
  '8–10: 223 (33,4 %)')
P('Tổng GPA: Dưới 5: 10 (1,5 %) | 5–6,9: 108 (16,2 %) | 7–7,9: 229 (34,3 %) | '
  '8–10: 315 (47,2 %)')

H('Bắt nạt mạng (Cyberbullying)', level=2)
P('"Bạn đã từng bị bắt nạt mạng chưa?"')
P('Chưa bao giờ: 316 (47,3 %) | Hiếm khi: 200 (29,9 %) | Đôi khi: 135 '
  '(20,2 %) | Thường xuyên: 14 (2,1 %)')
P('Bắt nạt mạng và cấp học: THPT có khả năng trải qua bắt nạt mạng đáng kể '
  'hơn (M tần suất = 3,35 vs THCS 2,35) và đau khổ liên quan đáng kể hơn '
  '(M = 19,13 vs 14,90), p < 0,01.')

# --- Trang 102-103 ---
PageMark('--- Trang 102–103, UNICEF Việt Nam, 2022 ---')

H('Kết quả khảo sát giáo viên (n = 66)', level=2)

P('Người tham gia: 66 giáo viên. Hà Nội 12, Điện Biên 18, Đồng Tháp 17, Gia '
  'Lai 18. THCS 40 (60,6 %), THPT 24 (36,4 %).')

P('Tuổi trung bình: 39,61 (range 25–55). Giới tính: Nam 10 (15,2 %), Nữ 56 '
  '(84,8 %). Dân tộc: Kinh 54 (81,8 %), Thái 3 (4,5 %), Thổ 1 (1,5 %).')

P('Trình độ cao nhất đạt được: Cao đẳng 5 (7,6 %), Cao đẳng + đào tạo bổ '
  'sung 48 (72,7 %), Thạc sĩ 13 (19,7 %).')

P('Số năm kinh nghiệm giảng dạy: Trung bình 16,86 (range 1–29).')

P('Chuyên môn giảng dạy: Khoa học tự nhiên 17 | Toán 12 | Ngữ Văn 10 | Lịch '
  'sử 7 | Ngoại ngữ 9 | Khoa học xã hội khác 6 | Môn bổ trợ 6')

P('Sĩ số lớp trung bình: Hầu hết giáo viên có 30–40 HS/lớp; 38 % có 40–50 '
  'HS/lớp.')

P('1. Mức độ quan ngại về stress và hạnh phúc của học sinh?', bold=True)
P('Hoàn toàn không: 1 (1,5 %) | Một chút: 4 (6,1 %) | Phần nào: 25 (37,9 %) '
  '| Rất quan ngại: 35 (53 %)')
P('→ 91 % giáo viên lo ngại.')

P('2. Mức độ quan ngại về sức khoẻ tâm thần của học sinh?', bold=True)
P('Hoàn toàn không: 0 | Một chút: 3 (4,5 %) | Phần nào: 27 (40,9 %) | Rất '
  'quan ngại: 36 (54,5 %)')
P('→ 95 % giáo viên lo ngại.')

P('5. Các vấn đề sức khoẻ tâm thần quan sát được ở học sinh:', bold=True)
P('• Chú ý kém 61 (92,4 %) | Thiếu tự tin 46 (69,7 %) | Năng lượng thấp/mệt '
  '40 (60,6 %) | Lo âu 34 (51,5 %) | Tức giận/hung hăng 12 (18,2 %) | Buồn/'
  'trầm cảm 10 (15,2 %)')

P('3. Tỷ lệ học sinh bị stress trong tháng qua (ước tính của giáo viên):', bold=True)
P('0–10 %: 34 GV | 10–20 %: 15 GV | 20–30 %: 3 GV | 30–40 %: 7 GV | 40–50 %: '
  '2 GV | 50–60 %: 2 GV | 60–70 %: 1 GV | 70–80 %: 2 GV | 80–90 %: 0 | 90–'
  '100 %: 0')

P('4. Tỷ lệ học sinh có vấn đề SKTT trong tháng qua (ước tính của giáo viên):',
  bold=True)
P('0–10 %: 50 GV | 10–20 %: 10 GV | 20–30 %: 3 GV | 30–40 %: 3 GV | > 40 %: 0')

P('5. Các vấn đề sức khoẻ tâm thần quan sát được ở học sinh (đã nêu ở trên, '
  'bổ sung đầy đủ):', bold=True)
P('Lo âu: 34 (51,5 %) | Chú ý kém: 61 (92,4 %) | Năng lượng thấp/mệt: 40 '
  '(60,6 %) | Buồn/trầm cảm: 10 (15,2 %) | Thiếu tự tin: 46 (69,7 %) | Tức '
  'giận/hung hăng: 12 (18,2 %) | Hành vi gây rối trong lớp: 13 (19,7 %) | '
  'Hành vi đối lập/thách thức: 7 (10,6 %) | Lạm dụng rượu/chất: 3 (4,5 %) | '
  'Xung đột bạn bè nghiêm trọng: 6 (9,1 %)')

P('6. Stress có ảnh hưởng đến thành tích học tập của học sinh không?', bold=True)
P('Stress cải thiện thành tích: 2 (3 %) | Stress không ảnh hưởng: 3 (4,5 %) '
  '| Stress tác động tiêu cực đến thành tích: 61 (92,4 %)')

P('7. Các yếu tố đóng góp vào các vấn đề sức khoẻ tâm thần học sinh '
  '(theo mức độ)?', bold=True)

# Bảng đầy đủ yếu tố nguy cơ
t = doc.add_table(rows=12, cols=5)
t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
t.autofit = False
set_grid(t, [4.5, 2.8, 2.8, 2.8, 2.8])
tbl_data = [
    ('Yếu tố', 'Không phải yếu tố', 'Yếu tố nhỏ', 'Yếu tố vừa', 'Yếu tố lớn'),
    ('Hành vi xấu', '10 (15,2 %)', '28 (42,4 %)', '18 (27,3 %)', '4 (6,1 %)'),
    ('Di truyền', '17 (25,8 %)', '25 (37,9 %)', '16 (24,2 %)', '6 (9,1 %)'),
    ('Vấn đề gia đình', '4 (6,1 %)', '6 (9,1 %)', '25 (37,9 %)', '28 (42,4 %)'),
    ('Áp lực học tập', '6 (9,1 %)', '20 (30,3 %)', '27 (40,9 %)', '11 (16,7 %)'),
    ('Stress cuộc sống', '4 (6,1 %)', '9 (13,6 %)', '36 (54,5 %)', '15 (22,7 %)'),
    ('Tính cách', '14 (21,2 %)', '23 (34,8 %)', '19 (28,8 %)', '5 (7,6 %)'),
    ('Môi trường', '6 (9,1 %)', '8 (12,1 %)', '30 (45,5 %)', '17 (25,8 %)'),
    ('Sang chấn', '9 (13,6 %)', '14 (21,2 %)', '31 (47 %)', '8 (12,1 %)'),
    ('Nuôi dạy con', '1 (1,5 %)', '7 (10,6 %)', '28 (42,4 %)', '26 (39,4 %)'),
    ('Nghiệp (Karma)', '48 (72,7 %)', '10 (15,2 %)', '5 (7,6 %)', '0'),
    ('Thất bại tâm linh', '45 (68,2 %)', '16 (24,2 %)', '1 (1,5 %)', '1 (1,5 %)'),
]
for ri, rd in enumerate(tbl_data):
    for ci, v in enumerate(rd):
        c = t.rows[ri].cells[ci]; c.text = str(v)
        colw(c, [4.5, 2.8, 2.8, 2.8, 2.8][ci])
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if ri == 0: r.bold = True
        if ri == 0: shade(c, 'D9E2F3')

P('Phát hiện then chốt: Gia đình (cả vấn đề và nuôi dạy) là yếu tố hàng đầu — '
  '79 % giáo viên coi là yếu tố vừa hoặc lớn. Stress cuộc sống (77 %), '
  'Môi trường (71 %) theo sau. Áp lực học tập 57 % yếu tố vừa/lớn. '
  'Ngược lại, các yếu tố "nghiệp" và "thất bại tâm linh" được hầu hết giáo '
  'viên xác định không phải yếu tố — cho thấy giáo viên có quan điểm khoa học '
  'về sức khoẻ tâm thần.', italic=True)

P('8. Khoảng bao nhiêu % trẻ em Việt Nam mắc vấn đề sức khoẻ tâm thần? '
  '(ước tính của giáo viên)', bold=True)
P('< 5 %: 13 (19,7 %) | 5–15 %: 29 (43,9 %) | 15–25 %: 17 (25,8 %) | 25–50 %: '
  '4 (6,1 %) | 50–75 %: 1 (1,5 %) | > 75 %: 0')
P('Lưu ý: hầu hết giáo viên đánh giá thấp tỷ lệ vấn đề sức khoẻ tâm thần ở '
  'trẻ em (trong khi tỷ lệ thực tế theo SDQ-25 là 26,1 %).', italic=True)

P('9. Nhận thức của giáo viên về năng lực cá nhân và vai trò hỗ trợ học sinh '
  'có vấn đề tâm lý – xã hội.', bold=True)
P('Hầu hết giáo viên cảm thấy tự tin về kiến thức cơ bản về sức khoẻ tâm '
  'thần học sinh, mặc dù một nhóm đáng kể (khoảng 36 %) cho rằng họ không '
  'có đủ thông tin về sức khoẻ tâm thần. Giáo viên cũng đồng ý phần lớn '
  'rằng việc xác định và hỗ trợ học sinh có các vấn đề sức khoẻ tâm thần '
  'là một phần vai trò của họ.')

P('10. Đào tạo về sức khoẻ tâm thần vị thành niên:', bold=True)
P('Hầu hết giáo viên (86,4 %) báo cáo rằng họ CHƯA NHẬN được đào tạo nào '
  'về sức khoẻ tâm thần vị thành niên — đây là khoảng trống nhân lực nghiêm '
  'trọng trong hệ thống giáo dục Việt Nam hiện tại.')

P('11. Rào cản chính đối với việc cung cấp thêm dịch vụ cho trẻ em có nhu '
  'cầu sức khoẻ tâm thần tại trường (theo đánh giá giáo viên):', bold=True)
P('• Thiếu nguồn nhân lực (tư vấn viên) — rào cản lớn nhất')
P('• Thiếu sự quan tâm/lo ngại từ phụ huynh')
P('• Thiếu cơ hội đào tạo cho nhân viên trường')
P('• Thiếu chính sách để hướng dẫn các dịch vụ')

P('12. Kiến thức của giáo viên về các chính sách hoặc chương trình cho sức '
  'khoẻ tâm thần học sinh:', bold=True)
P('• 81 % giáo viên báo cáo biết về các chính sách và/hoặc các chương trình '
  'cơ bản cho sức khoẻ tâm thần học sinh (bao gồm các cơ sở tư vấn học '
  'đường và nhân viên y tế trường học)')
P('• 58 % giáo viên nói rằng các nguồn lực và chương trình hiện có KHÔNG đủ '
  'để hỗ trợ học sinh')

# ============================================================
# TAI LIEU THAM KHAO — GIU NGUYEN TIENG ANH
# ============================================================
doc.add_page_break()
H('TÀI LIỆU THAM KHẢO (References)', level=1)
P('(Trang 108–117 của bản gốc. Theo chuẩn học thuật APA, danh sách tham khảo '
  'được GIỮ NGUYÊN TIẾNG ANH, không dịch.)', italic=True,
  color=RGBColor(0x66, 0x66, 0x66))

PageMark('--- Trang 108–117, UNICEF Việt Nam, 2022 ---')

# Load refs from parsed file
import importlib.util
spec = importlib.util.spec_from_file_location('refs_mod',
    'C:/Users/HLC/AppData/Local/Temp/unicef_vn/refs_list.py')
refs_mod = importlib.util.module_from_spec(spec)
spec.loader.exec_module(refs_mod)

for ref in refs_mod.REFS:
    P(ref, size=10, align='left')

# ============================================================
# PHAN REVIEW CUOI
# ============================================================
doc.add_page_break()
H('PHẢN BIỆN VÀ ĐÁNH GIÁ (Critique and Review) — Của người dịch', level=1, red=True)

Pred('BẢN DỊCH ĐẦY ĐỦ 118 TRANG UNICEF 2022 — HOÀN TẤT 12/04/2026', bold=True,
     italic=True)

Pred('Đánh giá chất lượng báo cáo UNICEF 2022:', bold=True)

Pred('ĐIỂM MẠNH:', bold=True)
Pred('• UNICEF + Bộ GD&ĐT — uy tín rất cao, báo cáo chính sách cấp quốc gia')
Pred('• Thiết kế mixed-methods chặt chẽ: khảo sát định lượng (668 HS + 66 GV) + '
     'FGD (39 HS + 21 phụ huynh + 35 GV) + KII (34 cán bộ)')
Pred('• 5 tỉnh đại diện: Hà Nội (Bắc, đô thị), Điện Biên (Bắc, DTTS, nông thôn), '
     'Gia Lai (Trung, DTTS), Đồng Tháp (Nam, nông thôn) — mất TPHCM do COVID-19')
Pred('• Công cụ đo chuẩn hoá: SDQ-25, MDS3 (Maryland School Climate), ESSA '
     '(Sun et al. 2011), RCBI-VN, đã validate tiếng Việt')
Pred('• Bao gồm các nhóm dễ tổn thương đặc biệt: DTTS (Mông, Gia-rai, Thái, Ba '
     'Na, Mường, Khơ Mú, Tày, Xơ đăng, v.v.) và LGBTQ')
Pred('• 9 khuyến nghị chính sách cho 4 nhóm đối tượng: Giáo dục, Y tế, Xã hội, '
     'UNICEF + khuyến nghị nghiên cứu thêm + hỗ trợ đại dịch')

Pred('HẠN CHẾ:', bold=True)
Pred('• Thiết kế cắt ngang — không xác lập nhân quả')
Pred('• Không bao gồm TPHCM (do COVID-19) → giảm tính khái quát hoá với các đô '
     'thị phía Nam')
Pred('• Báo cáo chính sách, KHÔNG phải nghiên cứu học thuật peer-reviewed '
     '(nhưng uy tín UNICEF bảo chứng)')
Pred('• n = 668 HS — cỡ mẫu vừa phải; nhỏ hơn Hoa 2024 [VN001] (n = 3.910) và '
     'Duong 2025 [VN029] (n = 2.631)')
Pred('• SDQ-25 đo vấn đề hành vi chung — KHÔNG đo lo âu hoặc trầm cảm riêng '
     '(khác GAD-7, PHQ-9, DASS-21). Điều này giới hạn khả năng đối chiếu trực '
     'tiếp với các nghiên cứu VN khác')
Pred('• Lấy mẫu thuận tiện (convenience sampling), không hoàn toàn ngẫu nhiên')
Pred('• Dữ liệu 2021 — có thể chưa phản ánh đầy đủ xu hướng hậu COVID, đặc biệt '
     'khi so với các nghiên cứu 2024-2025')
Pred('• 86,4 % giáo viên báo cáo chưa nhận đào tạo MHPSS — vấn đề nhân lực là '
     'rào cản hệ thống lớn nhất')

Pred('CÁC SỐ LIỆU THEN CHỐT (phục vụ đối chiếu liên bài — báo cáo Lo âu):', bold=True)
Pred('• 26,1 % học sinh VTN nguy cơ trung bình/cao vấn đề SKTT')
Pred('• 30,9 % vấn đề cảm xúc | 32 % vấn đề bạn bè | 14,4 % tăng động | 11 % '
     'vấn đề hành vi')
Pred('• Nữ > Nam: vấn đề cảm xúc (M = 4,92 vs 4,24, p < 0,01); áp lực học tập '
     '(M = 48,11 vs 44,27, p < 0,01)')
Pred('• THPT > THCS: vấn đề hành vi, tăng động, cảm xúc, tổng vấn đề (tất cả '
     'p < 0,01)')
Pred('• 28 % HS học > 3 giờ/đêm sau giờ học | 15 % HS có > 9 giờ học thêm/tuần')
Pred('• Cyberbullying: 52,2 % đã từng bị (hiếm/đôi khi/thường); THPT > THCS '
     'đáng kể')
Pred('• 91 % GV lo ngại stress HS | 95 % GV lo ngại SKTT HS | 86,4 % GV chưa '
     'nhận đào tạo MHPSS')
Pred('• 70 % trường ở tỉnh nông thôn thiếu phòng tư vấn riêng tư (Circular 31/'
     '2017 chưa thực hiện đầy đủ)')
Pred('• 668 HS × 5 tỉnh — đa dạng kinh tế xã hội, dân tộc, giới')

Pred('HƯỚNG NGHIÊN CỨU TIẾP THEO (Gap):', bold=True)
Pred('• RCT (thử nghiệm ngẫu nhiên có đối chứng) can thiệp MHPSS học đường — '
     'chưa có ngoài Happy House [VN030]')
Pred('• Đào tạo tư vấn viên học đường chuyên nghiệp — đầu tư nguồn nhân lực')
Pred('• Chính sách giảm áp lực học tập cấp quốc gia (MOET task force)')
Pred('• Mô hình hợp tác liên ngành MOET–MOH–MOLISA thí điểm — chưa có National '
     'Strategy on MH toàn diện')
Pred('• Nghiên cứu chuyên sâu về LGBTQ và DTTS — hiện mới chỉ có dữ liệu qualitative')

Pred('ĐÁNH GIÁ TỔNG THỂ: ⭐⭐⭐⭐⭐ (5/5)', bold=True)
Pred('Đây là nguồn tài liệu phong phú và có uy tín NHẤT về sức khoẻ tâm thần '
     'học sinh vị thành niên Việt Nam hiện có. Báo cáo cung cấp đồng thời: '
     '(a) dữ liệu định lượng từ 668 HS với công cụ chuẩn hoá; (b) insight định '
     'tính phong phú từ các bên liên quan; (c) phân tích chi tiết các chính '
     'sách MOET, MOH, MOLISA; (d) khuyến nghị hành động cụ thể cho mỗi ngành. '
     'Đây là tài liệu CỐT LÕI cho mọi đề cương nghiên cứu can thiệp MHPSS '
     'học đường tại Việt Nam.')

P('---', align='center', italic=True)
P('KẾT THÚC BẢN DỊCH ĐẦY ĐỦ — 10 BATCHES — 118 TRANG — UNICEF Việt Nam 2022',
  align='center', italic=True, size=10, bold=True)
P('Người dịch: Nhóm dự án Lo âu | Hoàn tất: 12/04/2026 | Format: Jenkins CTH v5 '
  '(QR, trang gốc, bảng Grid, phản biện đỏ)', align='center', italic=True, size=9)

# ============ save ============
doc.save(DST)
d2 = Document(DST)
print(f'Saved: {os.path.basename(DST)}')
print(f'Total: {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables')
total_chars = sum(len(p.text) for p in d2.paragraphs)
print(f'Total chars: {total_chars:,}')
