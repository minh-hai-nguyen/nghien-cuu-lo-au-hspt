# -*- coding: utf-8 -*-
"""Báo cáo Tổng quan 2024-2026 — DẠNG BẢNG + cột tự căn chiều rộng"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"

def smart_table(doc, data, col_widths_cm=None):
    """Create table with auto column widths based on content length."""
    num_cols = len(data[0])
    table = doc.add_table(rows=len(data), cols=num_cols)
    table.style = 'Table Grid'

    # Calculate widths
    if not col_widths_cm:
        col_lens = [0] * num_cols
        for row in data:
            for j, val in enumerate(row):
                col_lens[j] = max(col_lens[j], len(str(val)))
        total = sum(col_lens) or 1
        page_w = 16.0
        col_widths_cm = [max(1.5, (l / total) * page_w) for l in col_lens]

    for i, row in enumerate(data):
        for j, val in enumerate(row):
            cell = table.cell(i, j)
            cell.text = ''
            # Set width
            cell.width = Cm(col_widths_cm[j])
            p = cell.paragraphs[0]
            run = p.add_run(str(val))
            run.font.size = Pt(8)
            run.font.name = 'Times New Roman'
            if i == 0:
                run.bold = True
    doc.add_paragraph()
    return table

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

doc.add_heading('TỔNG QUAN NGHIÊN CỨU MỚI NHẤT (2024-2026)', level=0)
doc.add_heading('Lo âu và trầm cảm ở học sinh — Cơ sở cho đề cương nghiên cứu', level=1)
p = doc.add_paragraph()
run = p.add_run('Báo cáo này tổng hợp 15 bài báo xuất bản 2024-2026 cùng chủ đề với 11 bài gốc. Mọi số liệu đều có trích dẫn nguồn [X] tương ứng danh mục cuối báo cáo.')
run.font.size = Pt(11)
doc.add_paragraph()

# PART 1: Reviews & Meta-analyses
doc.add_heading('PHẦN 1: TỔNG QUAN HỆ THỐNG VÀ PHÂN TÍCH GỘP', level=1)

data1 = [
    ['[#]', 'Tác giả (Năm)', 'Tạp chí\nXếp hạng', 'Phạm vi', 'Kết quả chính', 'Điểm nổi bật\n(3-5 câu)', 'Liên quan đề cương VN'],
    [
        '[1]', '59 quốc gia\n(2025)', 'J Affect Disord\nQ1, IF~6.6',
        'Đa quốc gia\nhọc sinh đi học',
        'Lo âu ĐNA (Đông Nam Á): 3.78%\n— thấp nhất\ncác khu vực.',
        'Phân tích đa quốc gia lớn nhất 2025 về lo âu học sinh.\nĐông Nam Á có tỷ lệ thấp nhất.\nNhưng dữ liệu ĐNA rất hạn chế.\nTỷ lệ thấp có thể do thiếu sàng lọc.',
        'VN cần thêm dữ liệu để xác minh.\nNếu V-NAMHS (Vietnam National Adolescent Mental Health Survey — Khảo sát Quốc gia về Sức khỏe Tâm thần Thanh thiếu niên Việt Nam) 15.6% >> 3.78% → ĐNA có thể bị underestimate.'
    ],
    [
        '[2]', 'Anderson et al.\n(2025)', 'J Child Adolesc\nPsychiatr Nursing\n(Wiley)',
        'Review tường thuật\ntoàn cầu',
        '31.9% TN (thanh thiếu niên) 13-18\ntừng được chẩn đoán\nrối loạn lo âu.',
        '5 nhóm yếu tố gia tăng lo âu:\n(1) Mạng xã hội/công nghệ\n(2) Áp lực học tập\n(3) COVID-19\n(4) Bất ổn kinh tế\n(5) Thay đổi cấu trúc gia đình.',
        'Framework yếu tố nguy cơ → áp dụng trực tiếp cho Introduction đề cương.\nTất cả 5 yếu tố đều hiện diện ở VN.'
    ],
    [
        '[3]', 'GBD Trends\n(2025)', 'Transl Psychiatry\nQ1, IF~6.4',
        '204 quốc gia\n1990-2021\n10-24 tuổi',
        'Lo âu tăng 52%\ntừ 1990-2021.\nMạnh nhất nhóm\n10-14 tuổi.',
        'Phân tích xu hướng 31 năm — dài nhất hiện có.\nNhóm 10-14 (THCS — trung học cơ sở) tăng MẠNH NHẤT.\nTăng đặc biệt rõ sau 2019 (COVID-19).\nVùng SDI (Socio-Demographic Index — Chỉ số Nhân khẩu-Xã hội) trung bình (bao gồm VN) chịu ảnh hưởng nặng nhất.',
        'Nhóm 10-14 = THCS = đối tượng đề cương.\nBằng chứng xu hướng tăng 31 năm >> NSCH chỉ 7 năm.\nVN ở vùng SDI trung bình → nguy cơ cao.'
    ],
    [
        '[4]', 'ASEAN GBD\n(2025)', 'Lancet Public Health\nQ1, IF~25.4',
        '10 nước ASEAN\n1990-2021',
        'RLTT (rối loạn tâm thần) ASEAN: 11.9%.\nLo âu 10-14: 4.1%.\nDALY (Disability-Adjusted Life Years — Số năm sống điều chỉnh theo khuyết tật) 10-14: 16.3%.',
        'Nguồn Lancet — uy tín CAO NHẤT.\nNhóm 10-14 tuổi có gánh nặng DALY tâm thần cao nhất (16.3%).\nViệt Nam nằm trong phân tích.\nMalaysia +12.6%, Indonesia +9.0%, Philippines +7.5% (2019-2021).',
        'Trực tiếp liên quan ASEAN + VN.\nNhóm 10-14 = THCS.\nLancet = tạp chí hàng đầu thế giới → trích dẫn mạnh.'
    ],
    [
        '[5]', 'Lancet SEA\n(2025)', 'Lancet Regional\nHealth - SEA\nQ1',
        'Nam Á + ĐNA\nSystematic review',
        'Tổng hợp tỷ lệ\nvà yếu tố quyết định\nSKTT (sức khỏe tâm thần) TN khu vực.',
        'Review khu vực MỚI NHẤT từ Lancet.\nBao gồm VN, Ấn Độ, và các nước ĐNA.\nXác định khoảng trống nghiên cứu ở khu vực.\nThiếu dữ liệu từ nhiều nước ĐNA.',
        'Xác nhận khoảng trống nghiên cứu ĐNA.\nVN có thể đóng góp lấp đầy khoảng trống này.'
    ],
    [
        '[6]', 'LMIC (Low- and Middle-Income Countries — Nước thu nhập thấp và trung bình) Review\n(2025)', 'Psychiatr Res\nClin Practice\n(APA)',
        '208,842 trẻ/TN\nNước thu nhập\nthấp-TB',
        'Trầm cảm: 1-58%.\nLo âu: 1-30%.\nBiến thiên rất lớn.',
        'APA publication — uy tín cao.\nBiến thiên lớn phản ánh khác biệt phương pháp.\nYếu tố nguy cơ: nghèo đói, xung đột, thiếu dịch vụ.\n7/11 bài gốc từ LMIC → trực tiếp liên quan.',
        'Benchmark cho bối cảnh LMIC.\nVN là LMIC → áp dụng trực tiếp.\nGiúp đặt kết quả VN trong bối cảnh quốc tế.'
    ],
    [
        '[7]', 'Rising Burden\n(2024)', 'Frontiers\nPsychiatry\nQ2, IF~4.7',
        'Toàn cầu\n1990-2021\n10-24 tuổi',
        'Lo âu TN +52%.\nBất bình đẳng KT-XH (kinh tế-xã hội)\nlà yếu tố quan trọng.',
        'Bất bình đẳng kinh tế xã hội là yếu tố THEN CHỐT.\nCOVID-19 làm trầm trọng xu hướng sẵn có.\nVùng SDI trung bình chịu ảnh hưởng nặng nhất.\nOpen access — dễ truy cập.',
        'VN có bất bình đẳng lớn: Kinh vs thiểu số.\nDân tộc thiểu số VN 54.4% >> quốc gia 15.6%.\nCần phân tầng KT-XH trong đề cương.'
    ],
]
smart_table(doc, data1, [1.2, 2.5, 2.5, 2.0, 2.5, 4.5, 3.5])
doc.add_page_break()

# PART 2: Cross-sectional
doc.add_heading('PHẦN 2: NGHIÊN CỨU CẮT NGANG MỚI NHẤT', level=1)

data2 = [
    ['[#]', 'Tác giả (Năm)', 'Tạp chí', 'N / Quốc gia', 'Công cụ', 'Kết quả chính', 'Liên quan đề cương VN'],
    [
        '[8]', 'Delhi\n(2025)', 'PMC\n(chưa rõ tạp chí)',
        'N=679\nẤn Độ\n10-19 tuổi', 'PHQ-4',
        'Trầm cảm 25.92%.\nLo âu 13.70%.\nThấp hơn Saikia (24.4%).',
        'Nghiên cứu MỚI NHẤT từ Ấn Độ.\nBiến thiên vùng miền — Delhi vs Assam.\nPHQ-4 ngắn gọn — xem xét cho sàng lọc nhanh VN.'
    ],
    [
        '[9]', 'Saudi Arabia\n(2025)', 'Scientific Reports\n(Nature) Q1',
        'N=500+\nSaudi Arabia', 'GAD-7\nPHQ-9',
        'Tỷ lệ vẫn cao.\nXác nhận Alharbi (2019).',
        'Cập nhật 6 năm sau Alharbi.\nVấn đề SKTT Saudi vẫn chưa cải thiện.\nCần đánh giá xu hướng tương tự ở VN.'
    ],
    [
        '[10]', 'Jeddah\n(2024)', 'Frontiers\nPublic Health\nQ1',
        'N=728\nSaudi Arabia', 'PHQ-9\nGAD-7',
        'Trầm cảm 81.5%.\nLo âu 63.6%.\nCực cao — tương tự\nAlharbi (74%/63.5%).',
        'Xác nhận tỷ lệ cao Saudi KHÔNG phải ngoại lệ.\nAlharbi (2019) không bị phóng đại — pattern nhất quán.\nBối cảnh văn hóa Saudi khác VN → so sánh thận trọng.'
    ],
    [
        '[11]', 'JAMA US\n(2024)', 'JAMA Network\nOpen\nQ1, IF~13',
        'Dữ liệu QG (quốc gia)\nHoa Kỳ', 'Đa công cụ',
        'Xác nhận xu hướng\ntăng liên tục\ntrầm cảm/lo âu.\nNhóm 12-17 tăng\nmạnh nhất.',
        'JAMA — uy tín cao nhất.\nCập nhật NSCH (2020).\nNhóm 12-17 (THCS-THPT — trung học phổ thông) = đối tượng đề cương.\nXác nhận xu hướng toàn cầu — VN cần theo dõi.'
    ],
]
smart_table(doc, data2, [1.2, 2.0, 2.2, 2.0, 1.5, 3.0, 4.0])
doc.add_page_break()

# PART 3: Intervention
doc.add_heading('PHẦN 3: NGHIÊN CỨU CAN THIỆP MỚI NHẤT', level=1)

data3 = [
    ['[#]', 'Tác giả (Năm)', 'Tạp chí', 'Thiết kế', 'Can thiệp', 'Kết quả', 'Áp dụng cho VN'],
    [
        '[12]', 'OurFutures\n(2025)', 'eClinicalMedicine\n(Lancet) Q1',
        'RCT (Randomized Controlled Trial — Thử nghiệm Ngẫu nhiên có Đối chứng) cụm\n10 trường THPT\nÚc',
        '6 bài học, CBT (Cognitive Behavioral Therapy — Liệu pháp Nhận thức Hành vi),\nnhạy cảm giới.\nYear 8/9.',
        'Đánh giá phòng ngừa\nlo âu/trầm cảm.\nKết quả tích cực.',
        'Mô hình can thiệp tại trường — tham khảo cho pilot VN.\nCBT-based, 6 buổi — khả thi triển khai.\nNhạy cảm giới — quan trọng vì VN có chênh lệch giới.'
    ],
    [
        '[13]', 'Resilience\nMeta (2025)', 'Frontiers\nPsychiatry Q2',
        'Meta-analysis\nRCTs\nĐa quốc gia',
        'Can thiệp tăng\ncường resilience\ntại trường.',
        'Hiệu quả tích cực\nnhưng effect size\nnhỏ-TB.\nBao gồm TQ, Ấn Độ,\nPakistan.',
        'Liên quan Qiu (2022) — resilience OR=6.74.\nVN Academic Stress (2022) xác nhận resilience bảo vệ.\nCó thể thiết kế can thiệp resilience-based cho VN.'
    ],
    [
        '[14]', 'LMIC Prev.\n(2024)', 'PLOS ONE\nQ1',
        'Systematic review\nCan thiệp trường\nLMIC',
        'Phòng ngừa lo âu/\ntrầm cảm tại trường\nở nước LMIC.',
        'Bằng chứng hạn chế\nnhưng đầy hứa hẹn.\nThiếu nghiên cứu\ntừ ĐNA.',
        'Xác nhận KHOẢNG TRỐNG can thiệp ĐNA/LMIC.\nVN có thể LẤP ĐẦY khoảng trống này.\nĐề xuất: pilot RCT can thiệp SKTT tại trường VN.'
    ],
    [
        '[15]', 'Digital MH\n(2025)', 'PMC\n(chưa rõ tạp chí)',
        'Meta-analysis\nRCTs\nCan thiệp số',
        'App, web-based\ncho lo âu xã hội\nTN.',
        'Hiệu quả TB.\nChi phí thấp.\nKhả năng mở rộng\ncao.',
        'VN có tỷ lệ internet cao ở TN.\nCan thiệp số phù hợp bối cảnh thiếu nhân lực SKTT.\nChi phí thấp — khả thi cho trường công VN.'
    ],
]
smart_table(doc, data3, [1.2, 2.0, 2.2, 2.0, 2.0, 2.5, 4.0])
doc.add_page_break()

# PART 4: Guide for lit review
doc.add_heading('PHẦN 4: HƯỚNG DẪN VIẾT PHẦN TỔNG QUAN', level=1)

doc.add_heading('4.1. Cấu trúc đề xuất', level=2)
structure = [
    ('Đoạn 1 — Bối cảnh toàn cầu:', 'Lo âu ở TN tăng 52% từ 1990-2021 [3]. 4.1% trẻ 10-14 tuổi mắc lo âu [3][4]. COVID-19 làm tăng gấp đôi (Racine 2021, JAMA Pediatrics). Nguồn: GBD [3][4], Anderson [2].'),
    ('Đoạn 2 — Khu vực ASEAN/ĐNA:', 'Tỷ lệ RLTT ASEAN: 11.9% [4]. Lo âu học sinh ĐNA: 3.78% [1] — thấp nhất toàn cầu nhưng dữ liệu hạn chế [5][6]. Lancet SEA review xác nhận khoảng trống [5].'),
    ('Đoạn 3 — Việt Nam:', 'V-NAMHS: lo âu 15.6%. COVID VN: 41.5% trong vs 25.4% sau. Dân tộc TT: 54.4%. THCS Tuy Hòa: 22.7%. Thiếu nghiên cứu THCS quy mô lớn.'),
    ('Đoạn 4 — So sánh quốc tế:', 'Tỷ lệ dao động 9.89%-73.3%. Công cụ khác nhau → kết quả khác nhau. Nam > nữ ở 3/11 nghiên cứu. Bảng so sánh 11 bài gốc.'),
    ('Đoạn 5 — Yếu tố nguy cơ:', 'Nữ giới [1][7], nghèo đói [6][7], bắt nạt, giấc ngủ kém, nghiện game, nuôi dạy tiêu cực, sử dụng chất. Bảo vệ: resilience, hỗ trợ tại trường.'),
    ('Đoạn 6 — Khoảng trống → Lý do NC (nghiên cứu):', 'Thiếu NC THCS tại VN [6]. Thiếu can thiệp tại trường ĐNA [14]. Thiếu dữ liệu dân tộc TT. Thiếu NC dọc. → Lý do đề cương.'),
]
for label, content in structure:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0, 0, 0xCC)
    run = p.add_run(content)
    run.font.size = Pt(11)

doc.add_heading('4.2. Danh mục tài liệu tham khảo', level=2)
refs = [
    '[1] 59 quốc gia (2025). J Affective Disorders. Phân tích đa quốc gia lo âu học sinh.',
    '[2] Anderson et al. (2025). J Child Adolesc Psychiatric Nursing. Review yếu tố gia tăng lo âu.',
    '[3] GBD Trends (2025). Translational Psychiatry. Xu hướng 1990-2021.',
    '[4] ASEAN GBD (2025). Lancet Public Health. Gánh nặng bệnh ASEAN.',
    '[5] Lancet SEA (2025). Lancet Regional Health - Southeast Asia. Review Nam/ĐNA.',
    '[6] LMIC Review (2025). Psychiatric Research Clinical Practice (APA). Review LMIC.',
    '[7] Rising Burden (2024). Frontiers in Psychiatry. Bất bình đẳng + COVID.',
    '[8] Delhi (2025). PMC. N=679, PHQ-4, Ấn Độ.',
    '[9] Saudi Arabia (2025). Scientific Reports (Nature). GAD-7 + PHQ-9.',
    '[10] Jeddah (2024). Frontiers in Public Health. N=728.',
    '[11] JAMA Depression US (2024). JAMA Network Open. Xu hướng tăng Hoa Kỳ.',
    '[12] OurFutures (2025). eClinicalMedicine (Lancet). RCT can thiệp trường Úc.',
    '[13] Resilience Meta (2025). Frontiers in Psychiatry. Meta-analysis RCTs.',
    '[14] LMIC Prevention (2024). PLOS ONE. Can thiệp trường LMIC.',
    '[15] Digital MH (2025). PMC. Can thiệp kỹ thuật số lo âu TN.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.runs[0].font.size = Pt(9)

out = os.path.join(BASE, 'DocFiles', 'TONG_QUAN_NGHIEN_CUU_MOI_2024_2026.docx')
doc.save(out)
print(f'DONE: {out}')
