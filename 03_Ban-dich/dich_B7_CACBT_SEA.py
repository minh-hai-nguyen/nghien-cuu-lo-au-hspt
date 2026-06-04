# -*- coding: utf-8 -*-
"""Dịch B7 — CA-CBT ĐNA 2024 — Praptomojati et al. — Asian J Psychiatry"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1016/j.ajp.2023.103896', size=10)

add_heading(doc, 'Tổng quan hệ thống về Liệu pháp Nhận thức\u2013Hành vi Thích ứng Văn hóa (CA-CBT) cho rối loạn lo âu tại Đông Nam Á', 1)
h = doc.add_paragraph()
r = h.add_run('A systematic review of Culturally Adapted Cognitive Behavioral Therapy (CA-CBT) for anxiety disorders in Southeast Asia')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'A systematic review of Culturally Adapted Cognitive Behavioral Therapy (CA-CBT) for anxiety disorders in Southeast Asia'),
    ('Tác giả', 'Ardian Praptomojati (1,2)*, Ajeng Viska Icanervilia (3,4,5), Maaike H. Nauta (1), Theo K. Bouman (1)'),
    ('Cơ quan', '(1) Tâm lý Lâm sàng, ĐH Groningen, Hà Lan; (2) ĐH Gadjah Mada, Indonesia; (3-5) Y tế, ĐH Groningen + ĐH Gadjah Mada'),
    ('Tạp chí', 'Asian Journal of Psychiatry (Q1, IF \u2248 12,0)'),
    ('Xuất bản', '2024, Vol. 92, 103896, 9 trang'),
    ('DOI', '10.1016/j.ajp.2023.103896'),
    ('PROSPERO', 'CRD42022336376'),
    ('Loại NC', 'Tổng quan hệ thống (PRISMA) \u2014 tổng hợp tường thuật'),
    ('Phạm vi', '11 nước ĐNA (ASEAN + Timor-Leste); 7 NC: Indonesia 3, Thái Lan 2, Singapore 1, Malaysia 1'),
    ('Truy cập', 'Open Access \u2014 CC BY 4.0'),
])
add_page_ref(doc, '1\u20139', 'Asian Journal of Psychiatry', 'Vol. 92, 2024')

# TOM TAT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Bối cảnh: Liệu pháp Nhận thức\u2013Hành vi (CBT) đã được chứng minh là điều trị hiệu quả cho rối loạn lo âu. Tuy nhiên, CBT vẫn chủ yếu sử dụng các khái niệm và cấu trúc bắt nguồn từ văn hóa phương Tây, và đa số nghiên cứu tập trung vào dân số phương Tây. Chưa rõ điều này chuyển đổi thế nào sang văn hóa phi phương Tây như Đông Nam Á.')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Tìm kiếm hệ thống PubMed, PsycINFO, Embase, CENTRAL, GARUDA, Google Scholar. Tiêu chuẩn: CBT thích ứng văn hóa (CA-CBT) cho rối loạn lo âu ở cộng đồng ĐNA. Phân tích tường thuật theo Khung Thích ứng Văn hóa Điều trị (CTAF): thành phần ngoại vi (peripheral) và thành phần cốt lõi (core).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('Kết quả: 7 NC được chọn (1 RCT, 3 bán thực nghiệm, 3 ca lâm sàng). 2 NC thích ứng đa thành phần. 2 NC sửa đổi thành phần CỐT LÕI bằng tích hợp giá trị địa phương vào quá trình tái cấu trúc CBT. 3 NC thích ứng thành phần NGOẠI VI: tài liệu, ngữ nghĩa, ví dụ văn hóa, cấu trúc phiên. 1 RCT cho thấy CA-CBT CẢI THIỆN tốt hơn điều trị thông thường (TAU). Nhưng KHÔNG THỂ xác lập mức độ vượt trội của CA-CBT so với CBT không thích ứng.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

# DANH GIA NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'Asian J Psychiatry Q1 IF \u2248 12 \u2014 tạp chí tâm thần châu Á UY TÍN NHẤT.',
    'TỔNG QUAN HỆ THỐNG DUY NHẤT về CA-CBT lo âu tại ĐNA \u2014 rất phù hợp đề tài.',
    'Phạm vi 11 nước ĐNA (bao gồm VN, mặc dù 0 NC tìm thấy từ VN).',
    'Khung CTAF phân loại rõ: thích ứng ngoại vi vs cốt lõi \u2014 hướng dẫn thực hành.',
    'Phát hiện quan trọng: CHỈ 7 NC về CA-CBT lo âu ở ĐNA \u2014 GAP cực lớn.',
    '0 NC từ VN, Cambodia, Laos, Myanmar, Brunei, Philippines, Timor-Leste.',
    'Open Access CC BY \u2014 tải miễn phí.',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Chỉ 7 NC \u2014 quá ít để rút kết luận mạnh. 1 RCT duy nhất.',
    'Tổng hợp tường thuật (không meta-analysis) do thiếu dữ liệu.',
    'Không tìm thấy NC nào từ VN \u2014 dù VN nằm trong phạm vi tìm kiếm.',
    'Bài phát hành 2024 nhưng tìm kiếm chỉ đến 6/2023.',
]:
    add_p(doc, f'\u2022 {b}')

# 1. GIOI THIEU
add_page_ref(doc, '1\u20132', 'Asian J Psychiatry', 'Vol. 92, 2024')
add_heading(doc, '1. GIỚI THIỆU', 2)
add_p(doc, 'Rối loạn lo âu thuộc nhóm rối loạn tâm thần phổ biến nhất, với gánh nặng bệnh tật và chi phí y tế cao nhất (Bandelow & Michaelis, 2015). WHO ước tính 60,05 triệu người (23% dân số) sống với rối loạn lo âu tại ĐNA (WHO, 2017). Tại Indonesia, khoảng 8 triệu người (3,3%) mắc rối loạn lo âu. Tuy nhiên, không phải tất cả người bệnh tâm thần ở ĐNA nhận được giúp đỡ chuyên nghiệp do dịch vụ SKTT và nhân lực hạn chế.')
add_p(doc, 'CBT \u2014 một trong các phương pháp điều trị tâm lý được nghiên cứu nhiều nhất, đã chứng minh hiệu quả tốt qua nhiều định dạng \u2014 chủ yếu sử dụng khái niệm bắt nguồn từ văn hóa phương Tây. Điều này đặt câu hỏi về hiệu quả khi áp dụng cho dân số châu Á do khác biệt văn hóa (Rathod et al., 2018).')
add_p(doc, 'Thích ứng văn hóa (cultural adaptation) được định nghĩa là "sửa đổi có hệ thống một phương pháp điều trị dựa trên bằng chứng để phù hợp với ngôn ngữ, văn hóa, và bối cảnh sao cho tương thích với mô hình văn hóa, ý nghĩa, và giá trị của thân chủ" (Bernal et al., 2009). Khung CTAF (Cultural Treatment Adaptation Framework) chia thành: (1) Thích ứng NGOẠI VI \u2014 làm cho điều trị dễ hiểu và chấp nhận được; (2) Thích ứng CỐT LÕI \u2014 sửa đổi thành phần chính điều trị trung gian thay đổi triệu chứng.')

# 2. PHUONG PHAP
add_page_ref(doc, '2\u20133', 'Asian J Psychiatry', 'Vol. 92, 2024')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)
add_p(doc, 'PRISMA + PROSPERO (CRD42022336376). Tìm kiếm 7/2022, cập nhật 6/2023: PubMed, PsycINFO, Embase, CENTRAL, GARUDA (Indonesia), Google Scholar. Tiêu chuẩn: CA-CBT cho rối loạn lo âu ở cộng đồng 11 nước ĐNA (bao gồm VN). PTSD được bao gồm (trước DSM-5 thuộc lo âu). Không giới hạn tuổi, giới. Đánh giá thiên lệch: RoB 2 (RCT), ROBINS-I (bán thực nghiệm), JBI (ca lâm sàng).')

# 3. KET QUA
add_page_ref(doc, '3\u20137', 'Asian J Psychiatry', 'Vol. 92, 2024')
add_heading(doc, '3. KẾT QUẢ', 2)
add_p(doc, '1.653 bản ghi \u2192 loại trùng \u2192 1.073 sàng lọc \u2192 22 toàn văn \u2192 7 NC đủ tiêu chuẩn.')

add_heading(doc, 'Bảng 1. Đặc điểm 7 NC về CA-CBT lo âu ở ĐNA', 3)
add_table(doc,
    ['Quốc gia', 'Số NC', 'Thiết kế', 'Rối loạn', 'Giai đoạn'],
    [['Indonesia', '3', '1 bán TN + 2 ca LS', 'PTSD, SAD, Panic', '2007\u20132022'],
     ['Thái Lan', '2', '1 RCT + 1 bán TN', 'Đa lo âu, PTSD', '2007\u20132019'],
     ['Singapore', '1', 'Bán thực nghiệm', 'SAD', '2017'],
     ['Malaysia', '1', 'Ca lâm sàng', 'Lo âu chung', '2022'],
     ['VN, Cambodia, Laos...', '0', '\u2014', '\u2014', 'KHÔNG TÌM THẤY']],
    widths=[3.0, 1.5, 3.0, 3.0, 2.5])

add_heading(doc, 'Bảng 2. Loại thích ứng văn hóa (CTAF)', 3)
add_table(doc,
    ['Loại thích ứng', 'Số NC', 'Ví dụ cụ thể'],
    [['Ngoại vi \u2014 Tài liệu + Ngữ nghĩa', '3', 'Dịch thuật, điều chỉnh ngôn ngữ, ví dụ địa phương'],
     ['Ngoại vi \u2014 Cấu trúc phiên', '1', 'Điều chỉnh số phiên, thời lượng'],
     ['Cốt lõi \u2014 Tái cấu trúc nhận thức', '2', 'Tích hợp GIÁ TRỊ ĐỊA PHƯƠNG vào quá trình thay đổi suy nghĩ'],
     ['Đa thành phần', '2', 'Kết hợp ngoại vi + cốt lõi'],
     ['Không rõ ràng', '3', 'Không mô tả chi tiết thích ứng']],
    widths=[4.0, 1.5, 7.0])
add_p(doc, 'Phát hiện quan trọng: 2 NC tích hợp giá trị ĐỊA PHƯƠNG (ví dụ: giá trị gia đình, tôn giáo, tôn trọng thứ bậc) vào CBT tái cấu trúc \u2014 đây là thích ứng CỐT LÕI, không chỉ dịch thuật.', size=9, italic=True)

add_p(doc, '1 RCT duy nhất (Thái Lan): CA-CBT cải thiện tốt hơn điều trị thông thường (TAU). Tuy nhiên, không đủ dữ liệu để so sánh CA-CBT vs CBT KHÔNG thích ứng \u2014 không biết thích ứng có THỰC SỰ vượt trội hay CBT gốc cũng hiệu quả tương đương.')

# 4. THAO LUAN
add_page_ref(doc, '7\u20139', 'Asian J Psychiatry', 'Vol. 92, 2024')
add_heading(doc, '4. THẢO LUẬN VÀ KẾT LUẬN', 2)
add_p(doc, 'CHỈ 7 NC về CA-CBT lo âu ở ĐNA \u2014 rất ít. Gợi ý nhu cầu NC lớn hơn, đặc biệt RCT so sánh CA-CBT vs CBT gốc. Các thành phần thích ứng quan trọng: dịch thuật ngôn ngữ, ví dụ văn hóa địa phương, tích hợp giá trị gia đình/tôn giáo. Tại ĐNA, giá trị gia đình, tôn trọng thứ bậc, tôn giáo đóng vai trò quan trọng \u2014 CBT phương Tây (nhấn mạnh cá nhân, tự lập) có thể không phù hợp.')
add_p(doc, '0 NC từ VN mặc dù VN nằm trong phạm vi \u2014 GAP CỰC LỚN. VN có 60 triệu thanh niên, hệ thống y tế + giáo dục đang phát triển \u2014 cần CA-CBT riêng cho bối cảnh VN.')
add_p(doc, 'Hàm ý cho đề tài: Khi đề xuất can thiệp CBT tại trường VN (đề cương giai đoạn 2), cần thích ứng văn hóa: (1) dịch thuật + ví dụ VN, (2) tích hợp giá trị gia đình VN (cha mẹ, thầy cô), (3) cấu trúc phiên phù hợp lịch học VN. Phù hợp UNICEF VN22 (HS không thoải mái tìm GV) và Ireland QT32 (OGA bảo vệ).')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
for ref in [
    'Praptomojati, A., et al. (2024). A systematic review of CA-CBT for anxiety disorders in Southeast Asia. Asian J Psychiatry, 92, 103896.',
    'Bernal, G., et al. (2009). Handbook of U.S. Latino Psychology: Developmental and Community-Based Perspectives.',
    'Chu, J. & Leino, A. (2017). Advancement in the maturing science of cultural adaptations of evidence-based interventions. J Consulting Clin Psych, 85(1), 45\u201357.',
    '(Xem đầy đủ 60+ TLTK trong bài gốc)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('CA-CBT', 'Culturally Adapted Cognitive Behavioral Therapy \u2014 Liệu pháp Nhận thức\u2013Hành vi Thích ứng Văn hóa'),
    ('CBT', 'Cognitive Behavioral Therapy \u2014 Liệu pháp Nhận thức\u2013Hành vi'),
    ('CTAF', 'Cultural Treatment Adaptation Framework \u2014 Khung Thích ứng Văn hóa Điều trị'),
    ('TAU', 'Treatment As Usual \u2014 Điều trị Thông thường'),
    ('RCT', 'Randomized Controlled Trial'),
    ('PTSD', 'Post-Traumatic Stress Disorder \u2014 Rối loạn Stress Sau Sang chấn'),
    ('SAD', 'Social Anxiety Disorder \u2014 Rối loạn Lo âu Xã hội'),
    ('ĐNA', 'Đông Nam Á (Southeast Asia)'),
    ('PRISMA', 'Preferred Reporting Items for Systematic Reviews and Meta-Analyses'),
    ('PROSPERO', 'Prospective Register of Systematic Reviews'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'Asian J Psychiatry Q1 IF \u2248 12 \u2014 tạp chí uy tín nhất về tâm thần châu Á.',
    'TỔNG QUAN HỆ THỐNG DUY NHẤT về CA-CBT lo âu ở ĐNA. PRISMA + PROSPERO.',
    'Phạm vi 11 nước ĐNA \u2014 bao gồm VN. Tìm kiếm cả GARUDA (Indonesia) + Google Scholar.',
    'Khung CTAF phân loại rõ ràng: ngoại vi vs cốt lõi \u2014 hướng dẫn thực hành.',
    'Phát hiện 0 NC từ VN \u2014 xác nhận GAP cực lớn cho đề tài.',
    'Open Access CC BY.',
]:
    add_red(doc, f'\u2022 {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'CHỈ 7 NC \u2014 quá ít. 1 RCT duy nhất (Thái Lan). Không đủ cho meta-analysis.',
    'Không thể so sánh CA-CBT vs CBT KHÔNG thích ứng \u2014 câu hỏi chính chưa trả lời.',
    '3/7 NC không mô tả chi tiết thích ứng \u2014 thiếu minh bạch.',
    'Tìm kiếm chỉ đến 6/2023 \u2014 có thể bỏ sót NC mới.',
    'Bao gồm PTSD (trước DSM-5) \u2014 có thể không phù hợp phân loại hiện tại.',
    'Không có NC nào với VTN/HS \u2014 tất cả 7 NC đều người lớn hoặc không rõ tuổi.',
]:
    add_red(doc, f'\u2022 {s}')

add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'VN CẦN phát triển CA-CBT riêng \u2014 thích ứng cho văn hóa VN (gia đình, thầy cô, tôn trọng thứ bậc, áp lực thi cử).',
    'RCT CA-CBT tại VN so sánh với CBT gốc phương Tây \u2014 để xác định thích ứng có vượt trội không.',
    'CA-CBT cho VTN/HS \u2014 không chỉ người lớn. BMC NMA QT29: CBT SUCRA 0,66 cho VTN nhưng chưa rõ CA-CBT.',
    'Tích hợp giá trị gia đình VN: cha mẹ tham gia CBT (parent-involved CBT) \u2014 phù hợp văn hóa. Islam QT31: cha mẹ AOR=0,75 bảo vệ.',
    'So sánh CA-CBT ĐNA: Indonesia vs Thái Lan vs VN \u2014 cùng ĐNA nhưng văn hóa khác.',
    'CBT qua app/online + thích ứng VN \u2014 kết hợp B2 JMIR (digital SAD) + B7 (CA-CBT).',
]:
    add_red(doc, f'\u2022 {s}')

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '41_CACBT_SEA_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
checks = ['60,05', 'CTAF', 'PRISMA', 'PROSPERO', '1.653', '1.073', '7 NC', 'Indonesia', 'Thái Lan', 'CA-CBT']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
