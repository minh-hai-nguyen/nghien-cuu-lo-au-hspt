# -*- coding: utf-8 -*-
"""Sinh danh muc cong trinh EN (2 references) cho LA NCS Cong Thi Hang.
Format: Western order byline (Hang Thi Cong) + APA 7th journal name in italics.
29/05/2026 - v3: bo watermark/decoration."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', 'DanhMucCongTrinh_EN_v1_29052026.docx')

d = Document()

s = d.styles['Normal']
s.font.name = 'Times New Roman'
s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.15

for sec in d.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.0)
    sec.right_margin = Cm(2.0)


def add_heading(text, size=14, bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=12):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    return p


def add_para(text, size=12, bold=False, italic=False, align=WD_ALIGN_PARAGRAPH.LEFT,
             space_after=6, left_indent=None):
    p = d.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_after = Pt(space_after)
    if left_indent is not None:
        p.paragraph_format.left_indent = Cm(left_indent)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p


def add_reference(number, parts):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.25)
    p.paragraph_format.first_line_indent = Cm(-1.25)
    p.paragraph_format.space_after = Pt(12)

    r0 = p.add_run(f'{number}. ')
    r0.font.name = 'Times New Roman'
    r0.font.size = Pt(12)
    r0.bold = True

    for text, italic in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
        r.italic = italic
    return p


# ============================================================
# CONTENT
# ============================================================

# Tieu de song ngu
add_heading('DANH MỤC CÔNG TRÌNH KHOA HỌC CỦA TÁC GIẢ', size=14, bold=True, space_after=4)
add_heading('LIÊN QUAN ĐẾN LUẬN ÁN', size=14, bold=True, space_after=10)

add_heading('LIST OF SCIENTIFIC WORKS OF THE AUTHOR', size=14, bold=True, space_after=4)
add_heading('RELATED TO THIS DISSERTATION', size=14, bold=True, space_after=24)

# Reference 1
add_reference(1, [
    ('Hang Thi Cong (2026, May). Family environment and generalized anxiety disorder '
     'among lower secondary school students. ', False),
    ('Vietnam Journal of Educational Sciences', True),
    (', [vol/issue/pages]. [code: 2026.05.14 — chờ thầy clarify]', False),
])

# Reference 2
add_reference(2, [
    ('Hang Thi Cong (2026, May). Measurement validity of DSM-5 anxiety disorder scales '
     'in a sample of lower secondary school students in Vietnam. ', False),
    ('Vietnam Journal of Psychology', True),
    (', [vol/issue/pages].', False),
])

# ============================================================
# GHI CHU FORMAT
# ============================================================
add_heading('GHI CHÚ ĐỊNH DẠNG (APA 7th)',
            size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=10)

# 1. Author
add_para('1. Tên tác giả — Western order (thống nhất NCS + thầy ĐMĐ):',
         size=11, bold=True, space_after=4)
add_para('• Byline: Hang Thi Cong', size=11, space_after=2, left_indent=0.5)
add_para('• In-text English: Cong et al. (2026)', size=11, space_after=2, left_indent=0.5)
add_para('• In-text Vietnamese: Công Thị Hằng và cộng sự (2026)', size=11, space_after=2, left_indent=0.5)
add_para('• APA reference list nếu convert: Cong, H. T. (2026).', size=11, space_after=10, left_indent=0.5)

# 2. Year + month
add_para('2. Năm + tháng "(2026, May)":', size=11, bold=True, space_after=4)
add_para('• Giữ "(2026, May)" theo đúng text gốc.',
         size=11, space_after=2, left_indent=0.5)
add_para('• APA 7th journal article chuẩn chỉ là "(2026)" — month dùng cho magazine / '
         'newspaper / conference / advance online publication.',
         size=11, space_after=2, left_indent=0.5)
add_para('• Sau khi có vol/issue/pages → có thể chuyển thành "(2026)" cho đúng APA journal style.',
         size=11, italic=True, space_after=10, left_indent=0.5)

# 3. Title sentence case
add_para('3. Tiêu đề bài báo — sentence case:',
         size=11, bold=True, space_after=4)
add_para('• Chỉ viết hoa chữ đầu + danh từ riêng (DSM-5, Vietnam).',
         size=11, space_after=2, left_indent=0.5)
add_para('• Đã đúng cho cả 2 câu.',
         size=11, space_after=10, left_indent=0.5)

# 4. Journal name
add_para('4. Tên tạp chí — Title Case + in nghiêng:',
         size=11, bold=True, space_after=4)
add_para('• Câu 1: ',
         size=11, space_after=2, left_indent=0.5)
p_j1 = d.add_paragraph()
p_j1.paragraph_format.left_indent = Cm(0.5)
p_j1.paragraph_format.space_after = Pt(2)
r_j1 = p_j1.add_run('   Vietnam Journal of Educational Sciences')
r_j1.font.name = 'Times New Roman'
r_j1.font.size = Pt(11)
r_j1.italic = True
add_para('• Câu 2: ',
         size=11, space_after=2, left_indent=0.5)
p_j2 = d.add_paragraph()
p_j2.paragraph_format.left_indent = Cm(0.5)
p_j2.paragraph_format.space_after = Pt(2)
r_j2 = p_j2.add_run('   Vietnam Journal of Psychology')
r_j2.font.name = 'Times New Roman'
r_j2.font.size = Pt(11)
r_j2.italic = True
add_para('• Cần verify 2 tên official EN qua website tạp chí hoặc Ban biên tập.',
         size=11, italic=True, space_after=10, left_indent=0.5)

# 5. THCS naming
add_para('5. THCS — "lower secondary school" (MOET official):',
         size=11, bold=True, space_after=4)
add_para('• Đã apply thống nhất ở cả 2 câu để khớp LA chính.',
         size=11, space_after=10, left_indent=0.5)

# 6. Code
add_para('6. "code: 2026.05.14" — chờ clarify:',
         size=11, bold=True, space_after=4)
add_para('(a) Vol.Issue.Article (Vol 2026, số 5, bài 14)',
         size=11, space_after=2, left_indent=0.5)
add_para('(b) Ngày xuất bản 14/05/2026',
         size=11, space_after=2, left_indent=0.5)
add_para('(c) Manuscript ID nội bộ Tạp chí → sẽ bỏ khỏi reference',
         size=11, space_after=2, left_indent=0.5)
add_para('(d) DOI hoặc article number khác',
         size=11, space_after=10, left_indent=0.5)

# 7. Final form mau
add_heading('FORMAT FINAL DỰ KIẾN (sau khi đủ metadata)',
            size=12, bold=True, align=WD_ALIGN_PARAGRAPH.LEFT, space_after=8)

# Sample 1
p_s1 = d.add_paragraph()
p_s1.paragraph_format.left_indent = Cm(1.25)
p_s1.paragraph_format.first_line_indent = Cm(-1.25)
p_s1.paragraph_format.space_after = Pt(8)
r_s1a = p_s1.add_run('1. ')
r_s1a.font.name = 'Times New Roman'; r_s1a.font.size = Pt(11); r_s1a.bold = True
r_s1b = p_s1.add_run('Hang Thi Cong (2026). Family environment and generalized anxiety disorder '
                     'among lower secondary school students. ')
r_s1b.font.name = 'Times New Roman'; r_s1b.font.size = Pt(11)
r_s1c = p_s1.add_run('Vietnam Journal of Educational Sciences')
r_s1c.font.name = 'Times New Roman'; r_s1c.font.size = Pt(11); r_s1c.italic = True
r_s1d = p_s1.add_run(', 2026(5), 14–25. https://doi.org/10.xxxx/xxxxx')
r_s1d.font.name = 'Times New Roman'; r_s1d.font.size = Pt(11)

# Sample 2
p_s2 = d.add_paragraph()
p_s2.paragraph_format.left_indent = Cm(1.25)
p_s2.paragraph_format.first_line_indent = Cm(-1.25)
p_s2.paragraph_format.space_after = Pt(12)
r_s2a = p_s2.add_run('2. ')
r_s2a.font.name = 'Times New Roman'; r_s2a.font.size = Pt(11); r_s2a.bold = True
r_s2b = p_s2.add_run('Hang Thi Cong (2026). Measurement validity of DSM-5 anxiety disorder scales '
                     'in a sample of lower secondary school students in Vietnam. ')
r_s2b.font.name = 'Times New Roman'; r_s2b.font.size = Pt(11)
r_s2c = p_s2.add_run('Vietnam Journal of Psychology')
r_s2c.font.name = 'Times New Roman'; r_s2c.font.size = Pt(11); r_s2c.italic = True
r_s2d = p_s2.add_run(', X(Y), pp–pp. https://doi.org/10.xxxx/xxxxx')
r_s2d.font.name = 'Times New Roman'; r_s2d.font.size = Pt(11)

add_para('Dấu range trong số trang: en-dash "–" — không dùng hyphen "-".',
         size=10, italic=True, space_after=4)

# Clean metadata - bo het watermark
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
cp.content_status = ''
cp.identifier = ''
cp.language = ''
cp.revision = 1
cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'File size: {os.path.getsize(OUT)} bytes')
