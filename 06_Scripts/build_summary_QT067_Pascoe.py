"""Build Jenkins summary doc for QT067 Pascoe et al. 2020."""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/Tom-tat-tung-bai/QT067_Pascoe_AcademicStress_IJAY_2020.docx')

doc = Document()

# Margins
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

# Default font
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def heading(text, level=1, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    if level == 1:
        r.font.size = Pt(16)
    elif level == 2:
        r.font.size = Pt(14)
    else:
        r.font.size = Pt(12)
    if color:
        r.font.color.rgb = color
    return p

def para(text, bold=False, italic=False, color=None, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold; r.italic = italic
    r.font.size = Pt(size)
    if color:
        r.font.color.rgb = color
    return p

def bullet(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.italic = italic
    r.font.size = Pt(12)
    return p

# ===================== HEADER =====================
heading('QT067 — Pascoe, Hetrick & Parker (2020)', level=1, color=RGBColor(0x1F, 0x49, 0x7D))
heading('The impact of stress on students in secondary school and higher education', level=2)
para('International Journal of Adolescence and Youth, 25(1), 104–112.', italic=True)
para('DOI: 10.1080/02673843.2019.1596823', italic=True)
para('')

# ===================== METADATA TABLE =====================
heading('1. Thông tin định danh', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
meta = [
    ('Mã canonical', 'QT067'),
    ('Tác giả chính', 'Michaela C. Pascoe'),
    ('Đồng tác giả', 'Sarah E. Hetrick, Alexandra G. Parker'),
    ('Cơ quan tác giả chính', 'Institute for Health and Sport, Victoria University, Melbourne, Australia'),
    ('Năm xuất bản', '2020 (online 2019)'),
    ('Tạp chí', 'International Journal of Adolescence and Youth (IJAY)'),
    ('Loại bài', 'Narrative review (tổng quan tường thuật)'),
    ('Đối tượng tổng quan', 'Học sinh trung học (secondary) + sinh viên đại học (higher education)'),
    ('Khu vực', 'Quốc tế (chủ yếu USA, UK, Úc, Châu Á)'),
    ('DOI', '10.1080/02673843.2019.1596823'),
    ('Open Access', 'Có (CC BY 4.0; bản full text tại Victoria U repository)'),
    ('Email tác giả chính', 'michaela.pascoe@vu.edu.au'),
]
tbl = doc.add_table(rows=len(meta), cols=2)
tbl.style = 'Light Grid Accent 1'
for i, (k, v) in enumerate(meta):
    tbl.rows[i].cells[0].text = k
    tbl.rows[i].cells[1].text = v
    for cell in tbl.rows[i].cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11)
                r.font.name = 'Times New Roman'
        if cell == tbl.rows[i].cells[0]:
            for p in cell.paragraphs:
                for r in p.runs:
                    r.bold = True
para('')

# ===================== ABSTRACT =====================
heading('2. Tóm tắt (Abstract)', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
para('Nội dung tổng quan này khảo sát các tác động của stress học đường (academic stress / academic-related stress) lên học sinh trung học và sinh viên đại học theo 6 trục:')
bullet('Sức khỏe tâm thần (anxiety, depression).')
bullet('Sử dụng chất gây nghiện (rượu bia, ma túy, thuốc lá).')
bullet('Giấc ngủ (mất ngủ, chất lượng giảm).')
bullet('Sức khỏe thể chất (đau đầu, đau dạ dày, hệ miễn dịch).')
bullet('Thành tích học tập (giảm kết quả học, giảm khả năng tập trung).')
bullet('Tỷ lệ bỏ học (drop-out).')
para('Stress học đường không chỉ ảnh hưởng tâm lý mà còn lan tỏa sang hành vi sức khỏe và kết quả giáo dục. Nhóm tác giả đề xuất các chiến lược can thiệp toàn trường (whole-school) bao gồm chánh niệm (mindfulness), CBT, và cải thiện môi trường học.')
para('')

# ===================== KEY FINDINGS =====================
heading('3. Phát hiện chính', level=2, color=RGBColor(0x1F, 0x49, 0x7D))

heading('3.1. Sức khỏe tâm thần', level=3)
bullet('Stress học đường liên quan trực tiếp đến lo âu (anxiety) và trầm cảm (depression) ở thanh thiếu niên.')
bullet('Áp lực thi cử (high-stakes exams) là yếu tố nguy cơ độc lập với rối loạn lo âu.')
bullet('Sinh viên đại học báo cáo mức stress học đường cao tương đương stress công việc của người trưởng thành.')

heading('3.2. Sử dụng chất gây nghiện', level=3)
bullet('Học sinh stress cao có xu hướng dùng rượu/ma túy/thuốc lá để self-medicate.')
bullet('Liên kết hai chiều: stress → sử dụng chất, và sử dụng chất → stress tăng.')

heading('3.3. Giấc ngủ', level=3)
bullet('Stress làm giảm chất lượng giấc ngủ; thanh thiếu niên ngủ < 7h/đêm có nguy cơ trầm cảm cao gấp 2-3 lần.')
bullet('Mất ngủ vòng xoáy: thiếu ngủ → kém tập trung → kết quả học giảm → stress tăng.')

heading('3.4. Sức khỏe thể chất', level=3)
bullet('Triệu chứng cơ thể: đau đầu, đau dạ dày, mệt mỏi mãn tính.')
bullet('Stress mãn tính làm suy giảm hệ miễn dịch (cortisol kéo dài).')

heading('3.5. Thành tích học tập', level=3)
bullet('Stress quá mức làm giảm working memory và khả năng giải quyết vấn đề.')
bullet('Mối quan hệ hình chữ U ngược: stress thấp → động lực thấp; stress vừa → tối ưu; stress cao → giảm hiệu suất.')

heading('3.6. Bỏ học (drop-out)', level=3)
bullet('Stress không kiểm soát + hỗ trợ kém → tăng nguy cơ bỏ học.')
bullet('Sinh viên đại học bỏ học vì lý do tâm lý/stress chiếm tỷ lệ đáng kể trong nhóm bỏ học không hoàn cảnh.')

para('')

# ===================== METHODS =====================
heading('4. Phương pháp', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
para('Đây là TỔNG QUAN TƯỜNG THUẬT (narrative review), KHÔNG PHẢI hệ thống (systematic) hay meta-analysis:')
bullet('Không có PRISMA flow.')
bullet('Không có tiêu chí lựa chọn rõ ràng (no inclusion/exclusion criteria).')
bullet('Không có đánh giá chất lượng từng nghiên cứu (no risk-of-bias assessment).')
bullet('Tổng hợp định tính theo 6 chủ đề; không có effect size pooled.')
para('Dữ liệu trích dẫn từ ~80 nghiên cứu trước đó, phần lớn từ USA, UK, Úc.')
para('')

# ===================== STRENGTHS =====================
heading('5. Điểm mạnh', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
bullet('Khung tổng hợp 6-trục giúp người đọc hình dung tác động đa chiều của stress học đường.')
bullet('Open access (CC BY 4.0) — tiếp cận miễn phí, được trích dẫn rộng rãi (>1500 lần tính đến 2025).')
bullet('Ngôn ngữ dễ tiếp cận, phù hợp giáo viên, phụ huynh, người làm chính sách.')
bullet('Đề xuất can thiệp cụ thể (mindfulness, CBT, cải thiện môi trường học).')
para('')

# ===================== LIMITATIONS =====================
heading('6. Hạn chế', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
bullet('KHÔNG phải systematic review → không thể dùng làm "evidence summary" tin cậy cho meta-analysis.')
bullet('Thiếu dữ liệu Châu Á (chỉ vài trích dẫn Trung Quốc, Hàn Quốc, Nhật) — không có Việt Nam, Đông Nam Á.')
bullet('Không phân biệt rõ stress học đường vs stress chung; nhiều nghiên cứu trích dẫn dùng PSS (Perceived Stress Scale) chứ không phải thang đo academic stress chuyên biệt.')
bullet('Mối quan hệ nhân quả vs tương quan thường bị mờ trong narrative; người đọc cần thận trọng khi suy luận.')
bullet('Không cập nhật được các nghiên cứu sau 2018 (ngày submit).')
para('')

# ===================== APPLICATION VN =====================
heading('7. Liên hệ với đề tài Lo âu HS THCS/THPT Việt Nam', level=2, color=RGBColor(0xC0, 0x00, 0x00))

para('QT067 KHÔNG có dữ liệu Việt Nam, nhưng cung cấp framework tham khảo:', bold=True)
bullet('Khung 6-trục (mental health / substance use / sleep / physical health / achievement / drop-out) → áp dụng được khi viết bối cảnh trong báo cáo CTH; lưu ý rằng "substance use" ở VN có hình thái khác (game online, mạng xã hội thay vì rượu/ma túy như phương Tây).')
bullet('Mối quan hệ stress ↔ giấc ngủ → liên hệ với QT001 Jenkins (sleep hygiene), VN002 VNAMHS (rối loạn giấc ngủ trẻ VN).')
bullet('Mối quan hệ U-ngược giữa stress và performance → có thể giải thích vì sao một số HS THPT có stress cao nhưng vẫn học giỏi (cho tới ngưỡng tới hạn).')
bullet('Đề xuất whole-school intervention → liên kết với QT065 EACP (cấp lớp), QT066 peer support (cấp trường), CBT mobile app QT062 (cấp cá nhân).')

para('')
para('Cảnh báo khi trích dẫn:', bold=True, color=RGBColor(0xC0, 0x00, 0x00))
bullet('Không trích QT067 như một "evidence summary" — đây là narrative, không có effect size pooled.')
bullet('Trích để minh họa khung khái niệm hoặc làm rationale; số liệu cụ thể phải lấy từ nghiên cứu gốc (ví dụ Pascoe trích Caldwell 2019 — nếu cần dùng số, đọc Caldwell 2019 trực tiếp).')
bullet('Không dùng QT067 làm bằng chứng cho can thiệp ở VN — chỉ làm khung tham khảo. Cần kiểm chứng từ VN013 (Tô Thị Hồng), VN016 (Nguyễn Cao Minh), VNAMHS.')

para('')

# ===================== REFERENCES =====================
heading('8. Trích dẫn theo APA 7', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
para('Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823', italic=True)
para('')

# ===================== META-NOTE =====================
heading('9. Ghi chú quy trình', level=2, color=RGBColor(0x1F, 0x49, 0x7D))
para(f'Ngày tóm tắt: 2026-05-01.', italic=True, size=11)
para('Bản dịch chi tiết 1-1: 03_Ban-dich/QT067_Pascoe_AcademicStress_IJAY_2020.docx', italic=True, size=11)
para('PDF gốc: 02_Papers-goc/The-gioi_Au-My-Uc/QT067_Pascoe_AcademicStress_IJAY_2020.pdf', italic=True, size=11)
para('Áp dụng Nguyên tắc dịch v2 (10 nguyên tắc) — tham chiếu chéo (Nguyên tắc 9): xem critique trong bản dịch.', italic=True, size=11)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
