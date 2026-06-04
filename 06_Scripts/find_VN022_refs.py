# -*- coding: utf-8 -*-
import sys, io, os, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

files_to_check = []
for root, dirs, files in os.walk('.'):
    if '.git' in root or '_Archive' in root or '__pycache__' in root:
        continue
    for f in files:
        if f.endswith('.docx') and not f.startswith('~'):
            files_to_check.append(os.path.join(root, f))

print(f'Total docx files: {len(files_to_check)}')
print('='*70)

patterns = [
    ('249 GV', '249 teachers claim'),
    ('Kon Tum', 'Kon Tum province'),
    ('Ninh Thu[aạ]n', 'Ninh Thuan province'),
    (r'47\s*%.*học thêm|học thêm.*47\s*%|47\s*% HS học', '47% students extra classes'),
    ('UNICEF.*249', 'UNICEF with 249'),
    (r'\bVN(?:22|022)\b', 'VN22/VN022 canonical ID'),
]

for f in files_to_check:
    try:
        d = Document(f)
        txt = '\n'.join(p.text for p in d.paragraphs)
        for tbl in d.tables:
            for r in tbl.rows:
                txt += '\n' + ' | '.join(c.text for c in r.cells)
        hits = []
        for pat, desc in patterns:
            for m in re.finditer(pat, txt):
                ctx = txt[max(0,m.start()-60):m.end()+80].replace('\n', ' ')
                hits.append((desc, ctx[:200]))
                break
        if hits:
            short = os.path.basename(f)
            print(f'\n=== {short} ===')
            for desc, ctx in hits:
                print(f'  [{desc}]: ...{ctx}...')
    except Exception as e:
        pass

# Also check MD files
print('\n\n=== MD FILES ===')
md_files = []
for root, dirs, files in os.walk('.'):
    if '.git' in root or '_Archive' in root:
        continue
    for f in files:
        if f.endswith('.md'):
            md_files.append(os.path.join(root, f))

for f in md_files:
    try:
        with open(f, encoding='utf-8') as fh:
            txt = fh.read()
        hits = []
        for pat, desc in patterns:
            for m in re.finditer(pat, txt):
                ctx = txt[max(0,m.start()-60):m.end()+80].replace('\n', ' ')
                hits.append((desc, ctx[:200]))
                break
        if hits:
            short = os.path.basename(f)
            print(f'\n=== {short} ===')
            for desc, ctx in hits:
                print(f'  [{desc}]: ...{ctx}...')
    except Exception as e:
        pass
