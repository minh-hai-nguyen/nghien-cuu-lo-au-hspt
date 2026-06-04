# -*- coding: utf-8 -*-
"""Dịch đầy đủ QT32 — Ireland My World 2024 — Fitzgerald et al."""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link bài báo gốc: https://doi.org/10.1111/eip.13530', size=10)

# TIÊU ĐỀ
add_heading(doc, 'Khám phá xu hướng thay đổi trầm cảm và lo âu ở thanh thiếu niên từ 2012 đến 2019: Từ khảo sát My World cắt ngang lặp lại', 1)
h = doc.add_paragraph()
r = h.add_run('Exploring Changing Trends in Depression and Anxiety among Adolescents from 2012 to 2019: Insights from My World Repeated Cross-Sectional Surveys')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

# THÔNG TIN THƯ MỤC
add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Exploring Changing Trends in Depression and Anxiety among Adolescents from 2012 to 2019: Insights from My World Repeated Cross-Sectional Surveys'),
    ('Tiêu đề dịch', 'Khám phá xu hướng thay đổi trầm cảm và lo âu ở thanh thiếu niên từ 2012 đến 2019: Từ khảo sát My World cắt ngang lặp lại'),
    ('Tác giả', 'Amanda Fitzgerald (UCD Dublin), Ciara Mahon (UCD), Mark Shevlin (Ulster University), Barbara Dooley (UCD), Aileen O. Reilly (ALONE)'),
    ('Cơ quan', 'School of Psychology, University College Dublin; Ulster University, Coleraine; ALONE, Dublin'),
    ('Tạp chí', 'Early Intervention in Psychiatry (Q2, IF ≈ 2,5)'),
    ('Thông tin xuất bản', '2024, 11 trang'),
    ('DOI', '10.1111/eip.13530'),
    ('Loại NC', 'Cắt ngang lặp lại (repeated cross-sectional) — My World Survey'),
    ('Mẫu', 'Đợt 1 (2012): N = 5.490 VTN 12–19 tuổi; Đợt 2 (2019): N = 9.844 VTN 12–19 tuổi; Ireland'),
    ('Tài trợ', 'Jigsaw — Trung tâm Quốc gia về SKTT Thanh niên Ireland; ESB Energy for Generation Fund'),
])
add_page_ref(doc, '1–11', 'Early Intervention in Psychiatry', '2024')

# TÓM TẮT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Mục tiêu: Nghiên cứu chỉ ra tỷ lệ trầm cảm và lo âu ở VTN tăng trong ba thập kỷ qua. Tuy nhiên, các yếu tố cơ bản giải thích xu hướng tăng vẫn chưa được hiểu rõ. Nghiên cứu này kiểm tra các yếu tố nguy cơ và bảo vệ về tâm lý, xã hội và môi trường có thể giải thích thay đổi trong trầm cảm và lo âu ở VTN.')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Dữ liệu từ hai khảo sát My World đại diện quốc gia, VTN 12–19 tuổi, năm 2012 (N = 5.490) và 2019 (N = 9.844). Dữ liệu khảo sát về trầm cảm và lo âu (DASS-21) cùng 26 yếu tố nguy cơ/bảo vệ tiềm năng được đánh giá tại cả hai thời điểm. Phân tích đa nhóm (multiple group analysis) đánh giá liệu khả năng dự báo của các yếu tố có thay đổi từ Đợt 1 sang Đợt 2.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

p = doc.add_paragraph()
r = p.add_run('Kết quả: Tỷ lệ trầm cảm và lo âu TĂNG đáng kể giữa 2012 và 2019, đặc biệt ở nữ. Trầm cảm: M = 7,25 (2012) → 9,88 (2019), F = 246,60***. Lo âu: M = 6,37 → 9,41, F = 449,12***. Các yếu tố dự báo giải thích 37–61% phương sai. Một số yếu tố nhất quán qua cả hai đợt (bắt nạt, phân biệt đối xử, lạc quan), nhưng nữ giới, tìm kiếm trợ giúp chuyên môn, tự trọng thấp hơn, và khả năng phục hồi thấp hơn dự báo lo âu MẠNH HƠN ở Đợt 2 so với Đợt 1.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_p(doc, 'Kết luận: Phát hiện nhấn mạnh nhu cầu ưu tiên cung cấp dịch vụ SKTT VTN, đặc biệt cho nữ. Tự trọng và khả năng phục hồi là mục tiêu can thiệp quan trọng. Cần NC thêm để hiểu yếu tố nhân quả liên quan đến tăng lo âu và trầm cảm.')

# TÓM TẮT ĐÁNH GIÁ NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'My World Survey — khảo sát SKTT thanh niên lớn nhất Ireland, 2 đợt (2012 vs 2019).',
    'Mẫu cực lớn: 5.490 + 9.844 = 15.334 VTN 12–19 tuổi.',
    'Lo âu TĂNG 48% (M: 6,37 → 9,41); Trầm cảm TĂNG 36% (M: 7,25 → 9,88).',
    'Xu hướng TRƯỚC COVID — xác nhận tăng không chỉ do đại dịch. Phù hợp Norway 2025 (QT21, 13 năm tăng), JAACAP 2024 (QT23, lo âu gấp đôi Mỹ), GBD 2025 (QT30, AAPC 0,84%).',
    'Multi-group path analysis — kiểm tra yếu tố nào THAY ĐỔI giữa 2 đợt (phương pháp nâng cao).',
    'One Good Adult (OGA) — yếu tố bảo vệ đặc sắc: có 1 người lớn tin tưởng giảm lo âu/trầm cảm.',
    'Tự trọng và resilience MẠNH HƠN ở 2019 — gợi ý các yếu tố này trở nên quan trọng hơn theo thời gian.',
    '26 yếu tố nguy cơ/bảo vệ được đánh giá đồng thời — toàn diện.',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Chỉ Ireland — giàu, Tây Âu, văn hóa khác VN/châu Á.',
    'Cắt ngang lặp lại — không theo dõi cùng cá nhân, không xác lập nhân quả.',
    'DASS-21 sàng lọc — không chẩn đoán. So V-NAMHS 2022 (DISC-5 chẩn đoán: 2,3%).',
    'Đợt 2 chỉ đến 2019 — không phản ánh hậu COVID.',
    'Mẫu 2019 lớn hơn nhiều (9.844 vs 5.490) — có thể ảnh hưởng so sánh.',
]:
    add_p(doc, f'• {b}')

add_p(doc, 'Hướng cải thiện:', bold=True)
for b in [
    'NC tương tự tại VN: khảo sát lặp lại (trước/sau COVID) với DASS-21.',
    'Áp dụng OGA cho VN — xây dựng chương trình mentor tại trường.',
    'NC dọc (longitudinal) để xác nhận nhân quả: tự trọng thấp → tăng lo âu hay ngược lại?',
]:
    add_p(doc, f'• {b}')

# 1. GIỚI THIỆU
add_page_ref(doc, '1–2', 'Early Intervention in Psychiatry', '2024')
add_heading(doc, '1. GIỚI THIỆU', 2)

add_p(doc, 'Tuổi VTN là giai đoạn phát triển quan trọng, có thể định hình khả năng, mức độ nghiêm trọng và diễn biến của các vấn đề SKTT (Kessler et al., 2007, 2012). Có bằng chứng rằng 62,5% rối loạn tâm thần xảy ra trước 25 tuổi (Solmi et al., 2022). Tuy nhiên, các vấn đề SKTT ở nhóm tuổi này thường không được phát hiện cho đến khi trưởng thành và nhiều người trẻ không nhận được hỗ trợ đầy đủ (Hetrick et al., 2018).')

add_p(doc, 'Bằng chứng cho thấy tỷ lệ các vấn đề SKTT, đặc biệt lo âu và trầm cảm, đang tăng ở VTN (Bor et al., 2014; Collishaw, 2015). Mặc dù bằng chứng từ nước thu nhập thấp hạn chế, các NC từ nước thu nhập cao bao gồm Mỹ (Keyes et al., 2019; Lu, 2019), New Zealand (Fleming et al., 2014), Anh (Patalay & Gage, 2019), và các nước châu Âu khác (Thorisdottir et al., 2017; Tørmoen et al., 2020; Von Soest et al., 2014) đều chỉ ra xu hướng tăng, mặc dù tỷ lệ khác nhau giữa các nước.')

add_p(doc, 'Một số yếu tố được đề xuất đóng góp vào tăng SKTT: thay đổi cấu trúc gia đình, phụ thuộc kéo dài vào cha mẹ (Patton et al., 2016); tăng screen time và MXH (Carli et al., 2014); bắt nạt mạng (Bottino et al., 2015); áp lực trường học (Sweeting et al., 2010); giấc ngủ kém (Keyes et al., 2015); lối sống ít vận động (Felez-Nobrega et al., 2020); sử dụng cần sa và lo ngại ngoại hình (Von Soest et al., 2014).')

add_p(doc, 'Khái niệm "One Good Adult" (OGA — Một Người Lớn Tốt): có ít nhất một người trưởng thành đáng tin cậy trong cuộc sống VTN — được chứng minh là yếu tố bảo vệ SKTT quan trọng (Dooley et al., 2015; Sterrett et al., 2011). Điều này đặc biệt phù hợp với văn hóa Việt Nam, nơi thầy/cô giáo và cha mẹ đóng vai trò quan trọng.')

# 2. PHƯƠNG PHÁP
add_page_ref(doc, '2–5', 'Early Intervention in Psychiatry', '2024')
add_heading(doc, '2. PHƯƠNG PHÁP', 2)

add_p(doc, '2.1. Mẫu và quy trình', bold=True)
add_p(doc, 'My World Survey: khảo sát cộng đồng lớn nhất Ireland về SKTT thanh niên. Đợt 1 (My World Survey 1; Dooley & Fitzgerald, 2013): n = 5.490 VTN 12–19 tuổi, năm 2012. Đợt 2 (My World Survey 2; Dooley et al., 2019): n = 9.844 VTN 12–19 tuổi, năm 2019. Phân tích hiện tại sử dụng mẫu cắt giảm: Đợt 1 n = 6.085, Đợt 2 n = 5.869 (sau loại trừ dữ liệu thiếu).')
add_p(doc, 'Thu thập: Trường học được chọn ngẫu nhiên từ danh sách Bộ Giáo dục Ireland (Primary và Post-primary). Tất cả trường được gửi thư mời + biểu mẫu đồng ý. Khảo sát giấy hoặc trực tuyến tại trường, có hoặc không có trợ lý nghiên cứu.')
add_p(doc, 'Phê duyệt đạo đức: UCD School of Psychology.')

add_p(doc, '2.2. Thang đo', bold=True)
add_p(doc, '26 biến nguy cơ và bảo vệ được đánh giá (Bảng 1 gốc). Biến phụ thuộc:')
add_p(doc, '• Trầm cảm + Lo âu: DASS-21 (Lovibond & Lovibond, 1995). Tiểu thang trầm cảm (7 mục, 0–42) và lo âu (7 mục, 0–42). Likert 4 điểm (0–3), tổng × 2. Ngưỡng: Bình thường (trầm cảm 0–9, lo âu 0–7), Nhẹ–Trung bình (10–20 / 8–19), Nặng/Rất nặng (21–42 / 15–42). α Đợt 1: 0,88 (trầm cảm), 0,80 (lo âu); Đợt 2: 0,91, 0,84.')
add_p(doc, '• Triệu chứng loạn thần: APSS (Kelleher et al., 2011), 3 mục.')
add_p(doc, '• Rượu: AUDIT (Saunders et al., 1993), 10 mục, 0–40.')
add_p(doc, '• Lạc quan: LOT-R (Carver & Scheier, 1985), 6 mục, 0–24. α: 0,89/0,74.')
add_p(doc, '• Hài lòng cuộc sống: BMSLSS-PTTB (Seligson et al., 2003), 6 mục, 6–42. α: 0,81/0,84.')
add_p(doc, '• Khả năng phục hồi (Resilience): READ (Hjemdal et al., 2007) — 3 tiểu thang: Năng lực cá nhân (8 mục), Gắn kết gia đình (6 mục), Năng lực xã hội (5 mục).')
add_p(doc, '• Kết nối bạn bè + trường: MAC (Karcher & Lee, 2002), 6+6 mục.')
add_p(doc, '• Hỗ trợ xã hội: MSPSS (Zimet et al., 1988), 12 mục, 12–84. α: 0,93/0,95.')
add_p(doc, '• One Good Adult (OGA): 1 mục — "Bạn có 1 người lớn tin tưởng trong cuộc sống không?"')
add_p(doc, '• Ứng phó: 2 loại (ứng phó tập trung vấn đề + ứng phó tránh né).')
add_p(doc, '• Các biến khác: bắt nạt, phân biệt đối xử, hình ảnh cơ thể, cần sa, cha mẹ SKTT, tìm kiếm giúp đỡ (4 cấp).')

add_p(doc, '2.3. Phân tích', bold=True)
add_p(doc, 'SPSS 26 (thống kê mô tả, ANOVA, chi-square) + Mplus 8.4 (multi-group path analysis). Nhóm = Đợt 1/2. Trầm cảm và lo âu được hồi quy đồng thời trên 26 biến dự báo. Mô hình cơ bản: ràng buộc tất cả hệ số bằng nhau giữa 2 đợt ("stability model"). Chỉ số sửa đổi (MI > 5) cho biết ràng buộc nào cần nới lỏng. Đánh giá mô hình: CFI, TLI, RMSEA, SRMR.')

# 3. KẾT QUẢ
add_page_ref(doc, '5–8', 'Early Intervention in Psychiatry', '2024')
add_heading(doc, '3. KẾT QUẢ', 2)

add_p(doc, '3.1. Thay đổi trầm cảm và lo âu 2012 → 2019', bold=True)

add_heading(doc, 'Bảng 1. Thay đổi trầm cảm và lo âu 2012 → 2019 (DASS-21)', 3)
add_table(doc,
    ['Chỉ số', 'Đợt 1 (2012) M (SD)', 'Đợt 2 (2019) M (SD)', 'F', 'Thay đổi'],
    [['Trầm cảm (DASS-21)', '7,25 (8,58)', '9,88 (10,44)', '246,60***', '↑ TĂNG 36%'],
     ['Lo âu (DASS-21)', '6,37 (7,33)', '9,41 (8,98)', '449,12***', '↑ TĂNG 48%'],
     ['Tự trọng (Rosenberg)', '28,67 (5,70)', '26,46 (5,26)', '562,39***', '↓ GIẢM 8%'],
     ['Rượu (AUDIT)', '3,91 (5,81)', '2,75 (4,88)', '172,30***', '↓ GIẢM 30%'],
     ['Hỗ trợ xã hội (MSPSS)', '62,29 (15,49)', '64,46 (5,26)', '82,36***', '↑ Tăng nhẹ'],
     ['Triệu chứng loạn thần', '0,65 (0,85)', '0,87 (0,91)', '220,56**', '↑ Tăng']],
    widths=[4.0, 3.5, 3.5, 2.0, 2.5])
add_p(doc, '*** p < 0,001; ** p < 0,01. Lo âu tăng NHANH HƠN trầm cảm (48% vs 36%).', size=9, italic=True)

add_p(doc, '3.2. Thay đổi yếu tố nguy cơ và bảo vệ', bold=True)

add_heading(doc, 'Bảng 2. Thay đổi yếu tố nguy cơ/bảo vệ 2012 → 2019', 3)
add_table(doc,
    ['Yếu tố', 'Đợt 1 (%)', 'Đợt 2 (%)', 'χ²', 'Hướng'],
    [['Cha mẹ SKTT/nghiện', '13,3%', '17,0%', '37,81***', '↑ Tăng'],
     ['Cần sa', '12,0%', '15,2%', '30,90***', '↑ Tăng'],
     ['Cần giúp đỡ chuyên môn', '9,1%', '17,1%', '443,83***', '↑ TĂNG GẤP ĐÔI'],
     ['Phân biệt đối xử', '10,6%', '13,6%', '29,55***', '↑ Tăng'],
     ['Bắt nạt', '44,6%', '34,0%', '40,16***', '↓ GIẢM (tích cực)'],
     ['Hình ảnh cơ thể — hài lòng', '48,0%', '43,1%', '69,40***', '↓ Giảm']],
    widths=[4.0, 2.0, 2.0, 2.5, 3.0])
add_p(doc, 'Bắt nạt GIẢM nhưng lo âu vẫn TĂNG — gợi ý yếu tố khác đóng góp (tự trọng giảm, cha mẹ SKTT tăng, cần giúp đỡ chuyên môn tăng gấp đôi).', size=9, italic=True)

add_p(doc, '3.3. Yếu tố dự báo nhất quán qua cả 2 đợt', bold=True)
add_p(doc, 'Dự báo TIÊU CỰC lo âu và trầm cảm (giảm): lạc quan, kết nối bạn bè, năng lực cá nhân (resilience), tự trọng, tìm kiếm giúp đỡ không chính thức, One Good Adult (OGA).')
add_p(doc, 'Dự báo TÍCH CỰC lo âu (tăng): bắt nạt, phân biệt đối xử (nhất quán cả 2 đợt).')
add_p(doc, 'Hài lòng cơ thể chỉ liên quan trầm cảm (không lo âu). Hỗ trợ xã hội chỉ liên quan yếu tích cực với trầm cảm Đợt 1.')

add_p(doc, '3.4. Yếu tố THAY ĐỔI giữa 2 đợt (multi-group analysis)', bold=True)

add_heading(doc, 'Bảng 3. Yếu tố dự báo THAY ĐỔI giữa 2 đợt', 3)
add_table(doc,
    ['Yếu tố', 'Biến kết quả', 'Thay đổi', 'Ý nghĩa'],
    [['Nữ giới', 'Lo âu', 'MẠNH HƠN ở Đợt 2', 'Khoảng cách giới mở rộng'],
     ['Tìm kiếm giúp đỡ chuyên môn', 'Lo âu', 'MẠNH HƠN ở Đợt 2', 'Phản ánh nhu cầu tăng'],
     ['Tự trọng thấp', 'Lo âu + Trầm cảm', 'MẠNH HƠN ở Đợt 2', 'Tự trọng quan trọng hơn theo thời gian'],
     ['Resilience thấp (năng lực cá nhân)', 'Lo âu + Trầm cảm', 'MẠNH HƠN ở Đợt 2', 'Resilience quan trọng hơn'],
     ['Rượu', 'Lo âu', 'YẾU HƠN ở Đợt 2', 'Giảm vai trò (VTN uống ít hơn)'],
     ['Rượu', 'Trầm cảm', 'YẾU HƠN ở Đợt 2', ''],
     ['Hỗ trợ xã hội', 'Trầm cảm', 'YẾU HƠN ở Đợt 2', '']],
    widths=[4.0, 2.5, 3.0, 4.0])
add_p(doc, 'Phát hiện then chốt: Tự trọng và resilience trở nên QUAN TRỌNG HƠN cho lo âu/trầm cảm ở 2019 so với 2012. Rượu trở nên ÍT quan trọng hơn.', size=9, italic=True)

# 4. THẢO LUẬN
add_page_ref(doc, '7–10', 'Early Intervention in Psychiatry', '2024')
add_heading(doc, '4. THẢO LUẬN', 2)

add_p(doc, 'Sử dụng dữ liệu từ hai đợt mẫu lớn dựa trên cộng đồng, nghiên cứu này điều tra xu hướng SKTT VTN và kiểm tra 26 yếu tố nguy cơ/bảo vệ liên quan đến thay đổi mức lo âu và trầm cảm. Phù hợp với xu hướng tăng SKTT VTN toàn cầu (Keyes et al., 2019; Lu, 2019; Tørmoen et al., 2020), tỷ lệ lo âu và trầm cảm cao hơn năm 2019 so với 2012, đặc biệt ở nữ và VTN lớn tuổi hơn.')

add_p(doc, 'Đáng khích lệ, VTN Đợt 2 cho thấy mức cao hơn về ứng phó tập trung vấn đề, hỗ trợ xã hội, kết nối bạn bè, chấp thuận cha mẹ và có One Good Adult, cùng với mức thấp hơn về bắt nạt, rượu và phê phán cha mẹ so với Đợt 1. Phát hiện này phù hợp với báo cáo HBSC (Gavin et al., 2020): giảm bắt nạt, giảm rượu, cải thiện giao tiếp cha mẹ ở VTN 10–18 tuổi từ 1998–2018.')

add_p(doc, 'Tuy nhiên, VTN Đợt 2 cũng có mức thấp hơn về tự trọng, resilience cá nhân và lạc quan, và nhiều khả năng báo cáo triệu chứng loạn thần, sử dụng cần sa, vấn đề cần giúp đỡ chuyên môn, và cha mẹ có khó khăn SKTT/nghiện.')

add_p(doc, 'Giới tính:', bold=True)
add_p(doc, 'Nữ giới dự báo lo âu MẠNH HƠN ở Đợt 2 so với Đợt 1. Phù hợp với: khó khăn SKTT tăng ở nữ VTN nhanh hơn nam (Potrebny et al., 2017; Wiens et al., 2020). Khoảng cách giới mở rộng — phù hợp Úc 2025 (QT25): flourishing nữ 36,3% vs nam 52,2%. Ireland 2024 (QT32) và 59 Countries 2025 (QT31, AOR = 1,51) đều xác nhận.')

add_p(doc, 'Tự trọng và Resilience:', bold=True)
add_p(doc, 'Tự trọng thấp và resilience cá nhân thấp liên quan MẠNH HƠN với lo âu và trầm cảm ở Đợt 2. Có thể giải thích một phần xu hướng tăng: nếu tự trọng trở nên quan trọng hơn nhưng đồng thời giảm (28,67 → 26,46), thì tác động tích lũy sẽ lớn. Tuy nhiên, chiều nhân quả chưa rõ: tự trọng thấp → tăng lo âu hay lo âu tăng → giảm tự trọng? Cần NC dọc.')

add_p(doc, 'One Good Adult (OGA):', bold=True)
add_p(doc, 'Có OGA nhất quán liên quan giảm lo âu và trầm cảm qua cả 2 đợt. Đây là yếu tố bảo vệ đáng chú ý — có thể là thầy/cô giáo, cha mẹ, huấn luyện viên, hay người cố vấn. Tại VN, xây dựng chương trình mentor/OGA tại trường học có thể là can thiệp khả thi và hiệu quả — phù hợp Wen 2020 (hỗ trợ SKTT trường OR = 0,562).')

add_p(doc, 'Rượu giảm vai trò:', bold=True)
add_p(doc, 'Rượu liên quan yếu hơn với lo âu/trầm cảm ở Đợt 2. Phù hợp xu hướng toàn cầu: VTN uống ít rượu hơn trong thập kỷ qua (Von Soest et al., 2014). Tuy nhiên, cần sa TĂNG (12,0% → 15,2%) — có thể thay thế rượu.')

add_p(doc, 'Hạn chế:', bold=True)
add_p(doc, '• Cắt ngang lặp lại — không theo dõi cùng cá nhân. Không xác lập nhân quả. Li 2025 (QT22) dùng thiết kế dọc nhưng cũng gặp khó khăn nhân quả.')
add_p(doc, '• DASS-21 sàng lọc — không chẩn đoán. Tỷ lệ có thể cao hơn chẩn đoán lâm sàng. So V-NAMHS 2022: sàng lọc vs DISC-5 chênh 17 lần.')
add_p(doc, '• Chỉ Ireland — giàu, Tây Âu. Văn hóa khác châu Á/VN. Tuy nhiên, xu hướng tăng phù hợp toàn cầu.')
add_p(doc, '• Mẫu Đợt 2 lớn hơn nhiều (9.844 vs 5.490) — có thể tạo bất đối xứng thống kê.')
add_p(doc, '• Tự báo cáo — có thể bị ảnh hưởng bởi tăng nhận thức SKTT (prevalence inflation hypothesis — Foulkes & Andrews, 2023).')
add_p(doc, '• Đợt 2 (2019) — trước COVID. Không phản ánh tình hình hiện tại. Korea 2024 (QT34): SKTT xấu đi sau COVID.')

# 5. KẾT LUẬN
add_heading(doc, '5. KẾT LUẬN', 2)
add_p(doc, 'Dữ liệu 15.334 VTN Ireland qua 2 đợt (2012 và 2019) cho thấy lo âu tăng 48% và trầm cảm tăng 36%, đặc biệt ở nữ — xu hướng TRƯỚC COVID. Tự trọng và resilience trở nên QUAN TRỌNG HƠN cho SKTT VTN theo thời gian, trong khi đồng thời giảm — tạo vòng xoáy tiêu cực. One Good Adult (OGA) là yếu tố bảo vệ nhất quán. Bắt nạt giảm nhưng lo âu vẫn tăng — gợi ý yếu tố khác (MXH, tự trọng, cha mẹ SKTT) đang đóng góp.')
add_p(doc, 'Hàm ý: Cần ưu tiên dịch vụ SKTT VTN, đặc biệt cho nữ. Can thiệp xây dựng tự trọng + resilience + OGA tại trường — mục tiêu can thiệp quan trọng nhất dựa trên bằng chứng này.')

# TÀI LIỆU THAM KHẢO
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
refs = [
    'Fitzgerald, A., Mahon, C., Shevlin, M., Dooley, B., & O\'Reilly, A. (2024). Exploring changing trends in depression and anxiety among adolescents from 2012 to 2019. Early Intervention in Psychiatry.',
    'Dooley, B., & Fitzgerald, A. (2013). My World Survey 1: National Study of Youth Mental Health in Ireland. Dublin: Headstrong/UCD.',
    'Dooley, B., O\'Connor, C., Fitzgerald, A., & O\'Reilly, A. (2019). My World Survey 2: National Study of Youth Mental Health in Ireland. Dublin: Jigsaw/UCD.',
    'Lovibond, P.F. & Lovibond, S.H. (1995). The structure of negative emotional states: Comparison of the DASS with the Beck Depression and Anxiety Inventories. Behaviour Research and Therapy, 33, 335–343.',
    'Keyes, K.M., Gary, D., O\'Malley, P.M., Hamilton, A., & Schulenberg, J. (2019). Recent increases in depressive symptoms among US adolescents. Journal of Adolescent Health, 65, 590–598.',
    'Von Soest, T., Wichstrøm, L., & Kvalem, I.L. (2014). The development of global and domain-specific self-esteem from age 13 to 31. Journal of Personality and Social Psychology, 110, 592–608.',
    'Solmi, M., et al. (2022). Age at onset of mental disorders worldwide: Large-scale meta-analysis. Molecular Psychiatry, 27, 281–295.',
    'Foulkes, L. & Andrews, J.L. (2023). Are mental health awareness efforts contributing to the rise in reported mental health problems? New Ideas in Psychology, 69, 101010.',
]
for ref in refs:
    add_p(doc, ref, size=10)
add_p(doc, '(Xem danh mục đầy đủ trong bài gốc — 50+ tài liệu tham khảo)', size=10, italic=True)

# BẢNG VIẾT TẮT
add_abbreviation_table(doc, [
    ('DASS-21', 'Depression Anxiety Stress Scale 21 — Thang Trầm cảm Lo âu Căng thẳng 21 mục'),
    ('LOT-R', 'Life Orientation Test-Revised — Trắc nghiệm Định hướng Cuộc sống Chỉnh sửa'),
    ('AUDIT', 'Alcohol Use Disorders Identification Test — Bài kiểm tra Nhận diện Rối loạn Sử dụng Rượu'),
    ('READ', 'Resilience Scale for Adolescence — Thang Khả năng Phục hồi cho VTN'),
    ('BMSLSS', 'Brief Multidimensional Students\' Life Satisfaction Scale'),
    ('MAC', 'Hemingway Measure of Adolescent Connectedness'),
    ('MSPSS', 'Multidimensional Scale of Perceived Social Support — Thang Hỗ trợ Xã hội Đa chiều'),
    ('OGA', 'One Good Adult — Một Người Lớn Tốt'),
    ('APSS', 'Adolescent Psychotic-Like Symptom Screener'),
    ('CFI', 'Comparative Fit Index'),
    ('RMSEA', 'Root Mean Square Error of Approximation'),
    ('MI', 'Modification Index — Chỉ số Sửa đổi'),
    ('HBSC', 'Health Behaviour in School-aged Children'),
    ('UCD', 'University College Dublin'),
    ('SKTT', 'Sức khỏe tâm thần'),
    ('VTN', 'Vị thành niên'),
    ('MXH', 'Mạng xã hội'),
])

# PHẢN BIỆN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')

add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'My World Survey — khảo sát SKTT thanh niên lớn nhất Ireland. 2 đợt, tổng 15.334 VTN.',
    'Multi-group path analysis — phương pháp nâng cao kiểm tra YẾU TỐ NÀO thay đổi, không chỉ mô tả xu hướng. Phù hợp Norway 2025 (QT21) dùng decomposition.',
    '26 yếu tố nguy cơ/bảo vệ — toàn diện nhất trong các NC của Đề tài. AJP 2024 (QT28) chỉ review CBT/SSRI.',
    'Xu hướng TRƯỚC COVID (2012→2019) — xác nhận tăng lo âu/trầm cảm không chỉ do đại dịch. Phù hợp Norway 2025 (2011–2024), JAACAP 2024 (2013–2021, AOR=2,17), GBD 2025 (AAPC 0,84%).',
    'OGA (One Good Adult) — khái niệm đặc sắc, yếu tố bảo vệ nhất quán. Phù hợp Wen 2020 (hỗ trợ trường OR=0,562), 59 Countries (QT31, cha mẹ kiểm tra bài AOR=0,75).',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'Chỉ Ireland — giàu, Tây Âu, dân số nhỏ (~5 triệu). Văn hóa khác VN rất lớn (tôn giáo, gia đình, giáo dục). Tuy nhiên, xu hướng tăng phù hợp toàn cầu.',
    'Cắt ngang lặp lại — KHÔNG theo dõi cùng cá nhân. Không xác lập nhân quả (tự trọng thấp → tăng lo âu hay ngược lại?). Li 2025 (QT22) dùng dọc nhưng cũng gặp vấn đề nhân quả.',
    'DASS-21 sàng lọc — không chẩn đoán. Tỷ lệ cao có thể phản ánh triệu chứng, không phải rối loạn. So V-NAMHS 2022 (DISC-5: 2,3% chẩn đoán) — chênh lớn. Cần so sánh trên cùng mẫu (Gap #2 cross-study).',
    'Mẫu Đợt 2 lớn hơn nhiều (9.844 vs 5.490) — bất đối xứng có thể ảnh hưởng phân tích. Path analysis nhạy cảm với cỡ mẫu.',
    'Đợt 2 chỉ đến 2019 — KHÔNG phản ánh hậu COVID. Korea 2024 (QT34): SKTT đảo chiều sau COVID. Cần đợt 3 (2022+).',
    'Tự báo cáo — có thể bị ảnh hưởng bởi tăng nhận thức SKTT (Foulkes & Andrews, 2023: prevalence inflation hypothesis). Cùng thang đo nhưng VTN 2019 có thể BÁO CÁO triệu chứng dễ dàng hơn 2012.',
    'Không đo screen time/MXH — yếu tố quan trọng theo Norway 2025 (QT21: giải thích xu hướng) và Nature 2025 (QT27: g=0,46).',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống nghiên cứu / Research Gap:', bold=True)
for s in [
    'VN CẦN khảo sát lặp lại tương tự My World — so sánh trước/sau COVID. V-NAMHS 2022 là thời điểm 1, cần thời điểm 2 (2025–2026).',
    'Áp dụng OGA (One Good Adult) cho VN — xây dựng chương trình mentor tại trường THCS/THPT. Phù hợp văn hóa VN (thầy/cô, cha mẹ). Can thiệp chi phí thấp, dễ triển khai.',
    'Tự trọng và resilience ở VTN VN — chưa có NC. Ireland cho thấy 2 yếu tố này QUAN TRỌNG HƠN theo thời gian. VN có xu hướng tương tự không?',
    'So sánh Ireland vs VN — cùng dùng DASS-21 (Bảo Quyên 2025, Thảo Vi 2025, An Giang 2025) → có thể so sánh trực tiếp một số chỉ số.',
    'Kết hợp yếu tố bảo vệ (OGA, lạc quan, peer connectedness) vào can thiệp trường: Thảo Vi 2025 (VN19) cho thấy lạc quan giảm lo âu (β gián tiếp = −0,24). BMC NMA 2025 (QT29): ACT + CBT + PE hiệu quả.',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '32_Ireland_MyWorld_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
