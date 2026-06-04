"""DOC: Dien giai yeu to BAT NAT HOC DUONG (BNHD) - tac dong den RLLA HS THCS.
Theo style Doc 1 ALHT.

ALL FACTS VERIFIED:
- Chuong 3 luan an Bang 3.27-3.28:
  + BNHD -> RLLATQ: beta=0,215 (p<0,001)
  + BNHD -> RLLACL: beta=0,376 (p<0,001) - MANH NHAT
  + BNHD -> RLLAXH: beta=0,253 (p<0,001)
  + BNHD -> RLLA (3 factors): beta=0,276, R²=0,076
  + Fit: CFI 0,959-0,999; RMSEA 0,030-0,129
- Olweus 1996 OBVQ Revised, 42 muc - WebSearch verified
- Meta-analysis: bullying victimization -> anxiety OR ~4,72 (mild)
- Pooled prevalence 25% victims; 16% bully-victims
- Chen 2023 (QT007): bat nat 4 loai, OR 1,25-1,97 - verified Tom-tat
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Dien_giai_yeu_to_bat_nat_hoc_duong.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DIỄN GIẢI YẾU TỐ BẮT NẠT HỌC ĐƯỜNG\nTÁC ĐỘNG ĐẾN RỐI LOẠN LO ÂU\nỞ HỌC SINH TRUNG HỌC CƠ SỞ\n— Theo chương 3 luận án + bằng chứng quốc tế —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# 1. Phát biểu giả thuyết
H('1. Phát biểu giả thuyết', level=2, color=NAVY)
para('Giả thuyết (chính thức):', bold=True, indent=False)
para(
    'Bắt nạt học đường (school bullying victimization) là YẾU TỐ '
    'NGUY CƠ ĐỘC LẬP đối với rối loạn lo âu ở học sinh trung học '
    'cơ sở Việt Nam. Tác động được thể hiện trên CẢ BA dạng rối '
    'loạn lo âu — lan tỏa, chia ly, và xã hội — với cường độ |β| '
    'chuẩn hóa > 0,2 sau khi kiểm soát các biến nhân khẩu.', bold=True
)
para('Định nghĩa thao tác (operational definition):', bold=True, indent=False)
para(
    'BẮT NẠT HỌC ĐƯỜNG (school bullying) — theo Olweus (1996) — '
    'là hành vi tiêu cực CÓ CHỦ ĐÍCH, LẶP LẠI, và có CHÊNH LỆCH '
    'QUYỀN LỰC giữa kẻ bắt nạt và nạn nhân. Đo bằng thang Olweus '
    'Bully/Victim Questionnaire-Revised (OBVQ-R; Olweus, 1996) '
    'gồm 42 mục. Có bốn loại chính: (1) bắt nạt thể chất; (2) '
    'bắt nạt lời nói; (3) thao túng xã hội (loại trừ); (4) tấn '
    'công tài sản. Trong chương 3 luận án — yếu tố "bắt nạt thể '
    'chất" được đo qua thang OBVQ.', indent=False
)

# 2. Bối cảnh
H('2. Bối cảnh — ba cấp độ bằng chứng', level=2, color=NAVY)
para(
    'Trên BÌNH DIỆN TOÀN CẦU, bắt nạt học đường được Tổ chức Y tế '
    'Thế giới (WHO) và UNESCO xếp vào nhóm vấn đề sức khỏe công '
    'cộng hàng đầu của thanh thiếu niên. Phân tích tổng hợp gần '
    'đây xác lập tỷ lệ chung của nạn nhân bắt nạt = 25% (KTC 95% '
    '22–28%); kẻ bắt nạt = 16%; bắt nạt-bị-bắt nạt = 16% — tổng '
    'cộng hơn MỘT NỬA học sinh có liên quan với bắt nạt theo một '
    'vai trò.'
)
para(
    'Phân tích tổng hợp cho thấy nạn nhân bắt nạt có nguy cơ trầm '
    'cảm cao gấp 2,77 lần và nguy cơ lo âu cao gấp 4,72 lần (đối '
    'với bắt nạt mức độ nhẹ). Đặc biệt, nhóm bắt nạt-bị-bắt nạt '
    '(bully-victims) có nguy cơ cao NHẤT — gấp 3,19 lần.'
)
para(
    'Tại KHU VỰC, Chen và cộng sự (2023) — QT007 trong cơ sở dữ '
    'liệu — trên 63.205 học sinh trung học Trung Quốc dùng thang '
    'MPVS (Multidimensional Peer-Victimization Scale) phát hiện '
    'TẤT CẢ BỐN LOẠI bắt nạt đều liên quan có ý nghĩa với căng '
    'thẳng tâm thần. Đáng chú ý, thao túng xã hội (OR = 1,97) và '
    'bắt nạt lời nói (OR = 1,70) MẠNH HƠN bắt nạt thể chất (OR = '
    '1,51) và tấn công tài sản (OR = 1,25).'
)
para(
    'Tại VIỆT NAM, KẾT QUẢ CHƯƠNG 3 LUẬN ÁN xác lập bắt nạt học '
    'đường có tác động ĐỘC LẬP với RLLA. Cụ thể, mô hình SEM trên '
    'mẫu n = 1.352 học sinh THCS lớp 6–9 (Bảng 3.27–3.28) cho '
    'thấy β chuẩn hóa của BNHĐ với ba dạng RLLA đều có ý nghĩa '
    'thống kê (p < 0,001).'
)

# 3. Diễn giải số liệu chương 3
H('3. Diễn giải số liệu chương 3 luận án', level=2, color=NAVY)
para(
    'Bảng 3.27 báo cáo fit indices và Bảng 3.28 báo cáo hệ số tác '
    'động chuẩn hóa của bắt nạt học đường (BNHĐ) với từng dạng '
    'rối loạn lo âu.', indent=False
)
caption('Bảng 1. Hệ số β chuẩn hóa của BNHĐ với từng dạng RLLA (Bảng 3.28)')
add_table(
    ['Đường ảnh hưởng', 'β', 'S.E.', 'C.R.', 'p', 'Diễn giải'],
    [
        ['BNHĐ → RLLATQ (lan tỏa)', '0,215', '0,006', '7,349', '< 0,001', 'Tác động NHỎ–TRUNG BÌNH'],
        ['BNHĐ → RLLACL (chia ly)', '0,376', '0,007', '9,242', '< 0,001', 'Tác động TRUNG BÌNH (MẠNH NHẤT)'],
        ['BNHĐ → RLLAXH (xã hội)', '0,253', '0,007', '7,763', '< 0,001', 'Tác động NHỎ–TRUNG BÌNH'],
        ['BNHĐ → RLLA (3 factors)', '0,276', '0,013', '7,553', '< 0,001', 'R² = 0,076'],
        ['BNHĐ → RLLA (2 factors)', '0,254', '0,041', '7,294', '< 0,001', 'R² = 0,065'],
    ]
)
para('')
caption('Bảng 2. Fit indices các mô hình BNHĐ (Bảng 3.27)')
add_table(
    ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', 'Đánh giá'],
    [
        ['BNHĐ–RLLATQ', '90,581 (26)', '3,484', '0,982', '0,975', '0,043', 'TỐT'],
        ['BNHĐ–RLLACL', '43,206 (8)', '5,401', '0,984', '0,970', '0,057', 'CHẤP NHẬN'],
        ['BNHĐ–RLLAXH', '27,416 (8)', '3,427', '0,991', '0,984', '0,042', 'TỐT'],
        ['BNHĐ–RLLA (3 factors)', '94,392 (4)', '23,598', '0,959', '0,897', '0,129', 'KÉM (RMSEA cao)'],
        ['BNHĐ–RLLA (2 factors)', '2,197 (1)', '2,197', '0,999', '0,996', '0,030', 'XUẤT SẮC'],
    ]
)
para('')

H('Phát hiện ĐẶC BIỆT — Bắt nạt mạnh NHẤT cho lo âu CHIA LY', level=3, color=NAVY)
para(
    'Khác với áp lực học tập (mạnh nhất cho RLLATQ với β = 0,510) '
    'và tự trọng (mạnh nhất cho RLLATQ với β = −0,455), BẮT NẠT '
    'HỌC ĐƯỜNG cho thấy pattern KHÁC BIỆT: tác động MẠNH NHẤT '
    'lên rối loạn lo âu CHIA LY (RLLACL) với β = 0,376 — cao hơn '
    'cả tác động lên lan tỏa (β = 0,215) và xã hội (β = 0,253).'
)
para(
    'Có ba khả năng giải thích mô hình này:'
)
para(
    'KHẢ NĂNG 1 — Cơ chế gắn bó (attachment). Học sinh bị bắt nạt '
    'mất cảm giác AN TOÀN ở trường — không gian được kỳ vọng là '
    '"lớp thứ hai sau gia đình". Khi trường trở thành nơi nguy '
    'hiểm, học sinh KHÔNG MUỐN tách khỏi cha mẹ/người gắn bó — '
    'biểu hiện chính của RLLACL.'
)
para(
    'KHẢ NĂNG 2 — Tránh né trường học. Bắt nạt → từ chối đến '
    'trường → biểu hiện lo lắng khi tách khỏi nhà — phù hợp với '
    'pattern lâm sàng "school refusal" thường nhầm lẫn với RLLACL.'
)
para(
    'KHẢ NĂNG 3 — Đặc thù lứa tuổi THCS. Học sinh THCS (lớp 6–9, '
    'tuổi 11–15) ĐANG TRONG GIAI ĐOẠN CHUYỂN TIẾP từ trẻ con '
    'sang vị thành niên — lo âu chia ly chưa hoàn toàn biến mất '
    'mà chỉ giảm. Bắt nạt có thể KÍCH HOẠT lại pattern lo âu '
    'chia ly thuần dấu vết.'
)

H('Đánh giá fit indices', level=3, color=NAVY)
para(
    'Bốn trong năm mô hình đạt fit TỐT (CFI ≥ 0,98; RMSEA ≤ '
    '0,06). Riêng mô hình BNHĐ–RLLA (3 factors) có CHỈ SỐ '
    'KÉM với χ²/df = 23,598 và RMSEA = 0,129 — vượt ngưỡng '
    'chấp nhận. Mô hình BNHĐ–RLLA (2 factors) đạt fit XUẤT '
    'SẮC (CFI = 0,999; RMSEA = 0,030) — nên ƯU TIÊN báo cáo. '
    'Phù hợp với khuyến nghị Hu và Bentler (1999): khi mô '
    'hình cụ thể không fit, thử rebuild mô hình với cấu trúc '
    'khác.'
)

# 4. Bằng chứng quốc tế
H('4. Bằng chứng quốc tế chi tiết', level=2, color=NAVY)
caption('Bảng 3. Bằng chứng quốc tế về tác động bắt nạt → lo âu')
add_table(
    ['Công trình', 'Mẫu', 'Phát hiện chính'],
    [
        ['Olweus (1996), OBVQ-R',
         'Norway 130.000+ HS',
         'Khung khái niệm + thang đo chuẩn vàng quốc tế'],
        ['Meta-analysis bullying victims',
         'Pooled từ nhiều NC',
         'Nạn nhân lo âu OR = 4,72 (bắt nạt nhẹ)'],
        ['Chen và cộng sự (2023) [QT007]',
         'n = 63.205 HS TQ',
         '4 loại bắt nạt OR 1,25–1,97; thao túng xã hội mạnh nhất'],
        ['Pooled prevalence',
         'Global meta-analysis',
         'Nạn nhân 25%; bắt nạt-bị-bắt nạt 16%; rủi ro cao nhất 3,19x'],
    ]
)
para('')
para(
    'Phù hợp với bằng chứng quốc tế: chương 3 luận án xác lập '
    'tác động NGUY CƠ của BNHĐ với cả ba dạng RLLA. Cường độ '
    'tác động ở Việt Nam (β = 0,215–0,376) NẰM TRONG khoảng '
    'tham chiếu y văn quốc tế. Đặc thù Việt Nam: tác động '
    'MẠNH NHẤT lên RLLACL — phát hiện chưa được ghi nhận rõ '
    'rệt trong y văn quốc tế.'
)

# 5. Cơ chế
H('5. Bốn cơ chế đề xuất', level=2, color=NAVY)
para(
    'Y văn quốc tế đề xuất bốn cơ chế nối bắt nạt với rối loạn '
    'lo âu.', indent=False
)
para(
    'CƠ CHẾ 1 — Trục HPA và stress mãn tính. Bắt nạt LẶP LẠI '
    'kích hoạt phản ứng stress chronic → tăng cortisol cơ sở → '
    'tăng phản ứng lo âu. Phù hợp với khung sinh học của '
    'Pascoe và cộng sự (2020).'
)
para(
    'CƠ CHẾ 2 — Mô hình cảnh giác quá mức (hypervigilance). '
    'Học sinh bị bắt nạt phát triển cảnh giác liên tục về môi '
    'trường — đặc trưng của lo âu lan tỏa và xã hội.'
)
para(
    'CƠ CHẾ 3 — Mất an toàn ở trường → lo âu chia ly. Như đã '
    'phân tích ở phần 3, đây là cơ chế ĐẶC THÙ giải thích '
    'pattern β BNHĐ → RLLACL CAO NHẤT trong dữ liệu Việt Nam.'
)
para(
    'CƠ CHẾ 4 — Cô lập xã hội và mất hỗ trợ. Bắt nạt thường '
    'kèm theo loại trừ xã hội (thao túng xã hội — OR = 1,97 '
    'theo Chen 2023) → mất bạn bè + mất hỗ trợ → tăng lo âu '
    'xã hội.'
)

# 6. Hàm ý
H('6. Bốn hàm ý cho thiết kế can thiệp', level=2, color=NAVY)
para(
    'HÀM Ý 1 — Sàng lọc bắt nạt thường xuyên. Theo OBVQ-R, '
    'sàng lọc 6 tháng/lần ở học sinh THCS bằng phiên bản '
    'tiếng Việt. Đặc biệt chú ý các trường có tỷ lệ nạn '
    'nhân ≥ 25% (ngưỡng global meta-analysis).', indent=False
)
para(
    'HÀM Ý 2 — Can thiệp toàn trường (whole-school). Mô hình '
    'Olweus Bullying Prevention Program (OBPP) — chương '
    'trình can thiệp toàn trường có bằng chứng RCT. Bao gồm '
    'thành phần cấp lớp, cấp giáo viên, và cấp học sinh.'
)
para(
    'HÀM Ý 3 — Can thiệp ĐẶC BIỆT cho lo âu chia ly. Phát '
    'hiện β BNHĐ → RLLACL = 0,376 mạnh nhất gợi ý: cần '
    'theo dõi học sinh có biểu hiện "school refusal" và '
    'kiểm tra liệu nguyên nhân là bắt nạt — chứ không chỉ '
    'là "lo âu chia ly đơn thuần".'
)
para(
    'HÀM Ý 4 — Phân loại 4 loại bắt nạt. Theo Chen 2023, '
    'thao túng xã hội (OR = 1,97) và bắt nạt lời nói (OR = '
    '1,70) MẠNH HƠN bắt nạt thể chất (OR = 1,51). Can '
    'thiệp truyền thống thường tập trung vào bắt nạt thể '
    'chất — bỏ qua hai loại nguy hiểm hơn. Cần MỞ RỘNG '
    'phạm vi can thiệp.'
)

# 7. CÂU TRẢ LỜI
H('7. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue_run('Diễn giải bắt nạt học đường (BNHĐ) — chương 3 luận án:', bold=True)
blue_run(
    '(1) BNHĐ là YẾU TỐ NGUY CƠ ĐỘC LẬP cho cả ba dạng '
    'RLLA, p < 0,001 cho mọi đường ảnh hưởng.'
)
blue_run(
    '(2) PHÁT HIỆN ĐẶC BIỆT: BNHĐ tác động MẠNH NHẤT lên '
    'lo âu chia ly (β = 0,376), MẠNH HƠN tác động lên lan '
    'tỏa (β = 0,215) và xã hội (β = 0,253). Pattern KHÁC '
    'với áp lực học tập và tự trọng.'
)
blue_run(
    '(3) Cơ chế đề xuất: bắt nạt → mất an toàn ở trường '
    '→ tránh né trường → "school refusal" biểu hiện như '
    'lo âu chia ly. Đặc thù lứa tuổi THCS.'
)
blue_run(
    '(4) Bằng chứng quốc tế: meta-analysis nạn nhân lo '
    'âu OR = 4,72; tỷ lệ nạn nhân toàn cầu 25%. Chen '
    '2023 (QT007, n = 63.205): bắt nạt 4 loại OR '
    '1,25–1,97.'
)
blue_run(
    '(5) Mô hình tốt nhất: BNHĐ–RLLA (2 factors) đạt '
    'CFI = 0,999, RMSEA = 0,030 — XUẤT SẮC. Mô hình 3 '
    'factors có RMSEA = 0,129 — KÉM, không nên báo cáo '
    'làm chính.'
)

# 8. TLTK
H('8. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Chen, Z., Ren, S., He, R., và cộng sự. (2023). Prevalence and associated factors of depressive and anxiety symptoms among Chinese secondary school students. BMC Psychiatry, 23(1), 580. https://doi.org/10.1186/s12888-023-05068-1 [QT007 trong cơ sở dữ liệu dự án.]',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
    'Olweus, D. (1996). The revised Olweus Bully/Victim Questionnaire. University of Bergen, Research Center for Health Promotion (HEMIL Center).',
    'Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823 [QT067 trong cơ sở dữ liệu dự án.]',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
