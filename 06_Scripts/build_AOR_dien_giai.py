"""Build doc: AOR - Adjusted Odds Ratio - cach tinh."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

title = d.add_heading('AOR (Adjusted Odds Ratio — Tỷ số odds đã điều chỉnh): Cách tính và diễn giải', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Giải thích một trong những chỉ số phổ biến nhất trong nghiên cứu dịch tễ và tâm lý học')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# META
meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
mp = cell.paragraphs[0]
mp.add_run('Câu hỏi gốc của thầy: ').bold = True
mp.add_run('"AOR (Adjusted Odds Ratio) — tỷ số chênh lệch đã điều chỉnh được tính toán như thế nào?"')
mp2 = cell.add_paragraph()
mp2.add_run('Cách đọc tên: ').bold = True
mp2.add_run('AOR = Adjusted Odds Ratio (không phải "Ajusted"). Dịch sang tiếng Việt: "Tỷ số odds đã điều chỉnh" (một số sách dịch "Tỷ số chênh đã hiệu chỉnh"). Odds ≠ probability — hai khái niệm gần nhau nhưng không giống.')
d.add_paragraph()

# =========================================================
# PHẦN 1 — ÔN LẠI ODDS VÀ OR
# =========================================================
add_heading('1. Trước hết phải hiểu ODDS và OR thô (crude OR)', level=1)

add_para('1.1. Odds là gì?', bold=True, size=12)
add_para('Odds (tỷ số cược) = Xác suất xảy ra / Xác suất không xảy ra = P / (1 − P).')
add_para('Ví dụ: nếu 60 % HS lo âu → odds = 0,60 / 0,40 = 1,5. Nghĩa là "cứ 1 HS không lo âu thì có 1,5 HS bị lo âu" — cách diễn đạt của dân cá cược.')
d.add_paragraph()

# Bảng 2x2
add_para('1.2. OR thô (crude OR) từ bảng 2 × 2', bold=True, size=12)
add_para('Xét ví dụ: HS ngủ ít (<5h) có bị lo âu hay không (Zhu 2025 Trung Quốc, QT05)')
d.add_paragraph()

tb_tbl = d.add_table(rows=4, cols=4)
tb_tbl.style = 'Table Grid'
tbd = [
    ('', 'Lo âu (+)', 'Không lo âu (−)', 'Tổng'),
    ('Ngủ < 5h (phơi nhiễm +)', 'a = 180', 'b = 40', '220'),
    ('Ngủ ≥ 5h (phơi nhiễm −)', 'c = 400', 'd = 2.000', '2.400'),
    ('Tổng', '580', '2.040', '2.620'),
]
for i, row_data in enumerate(tbd):
    row = tb_tbl.rows[i]
    for j, v in enumerate(row_data):
        row.cells[j].text = v
        if i == 0 or j == 0:
            shade_cell(row.cells[j], 'D9E1F2')
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True
d.add_paragraph()

add_para('(Các con số minh hoạ — không phải số thật từ Zhu 2025)', color=(128, 128, 128))
d.add_paragraph()

add_para('Cách tính OR thô:')
for it in [
    'Odds lo âu trong nhóm ngủ ít = a/b = 180/40 = 4,5',
    'Odds lo âu trong nhóm ngủ đủ = c/d = 400/2.000 = 0,2',
    'OR thô = (a/b) / (c/d) = (a × d) / (b × c) = (180 × 2.000) / (40 × 400) = 22,5',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

add_para('→ OR thô = 22,5 → Odds lo âu của HS ngủ ít CAO GẤP 22,5 LẦN HS ngủ đủ.', bold=True, color=(192, 0, 0))
d.add_paragraph()

add_para('1.3. Vấn đề với OR thô — CONFOUNDING (nhiễu)', bold=True, size=12)
add_para('OR thô 22,5 có đáng tin không? KHÔNG, vì HS ngủ ít thường cũng: (a) áp lực học tập cao, (b) dùng điện thoại nhiều, (c) gia đình có xung đột, (d) lớp lớn hơn. Các yếu tố này ĐỒNG THỜI tác động lên lo âu → làm "phồng" (hoặc xì hơi) giá trị OR thô. Chúng gọi là CONFOUNDERS (biến gây nhiễu).')
d.add_paragraph()

add_para('Nếu chỉ báo cáo OR thô → khả năng sai lầm: "Ngủ ít gây lo âu". Thực tế có thể là "Áp lực học tập gây cả ngủ ít và lo âu". Giải pháp: ĐIỀU CHỈNH OR cho các biến đó → ra AOR.')
d.add_paragraph()

# =========================================================
# PHẦN 2 — AOR LÀ GÌ
# =========================================================
add_heading('2. AOR là gì?', level=1, color=(192, 80, 77))

add_para('AOR = OR đã được ĐIỀU CHỈNH cho các biến gây nhiễu khác (confounders). AOR ước tính "ảnh hưởng RIÊNG" của biến quan tâm, giữ cố định các biến khác.', bold=True, color=(192, 0, 0))
d.add_paragraph()

aor_tbl = d.add_table(rows=1, cols=1)
aor_tbl.style = 'Table Grid'
ac = aor_tbl.rows[0].cells[0]
shade_cell(ac, 'FFF8DC')
ap = ac.paragraphs[0]
ar = ap.add_run('Cách diễn đạt chuẩn: ')
ar.bold = True
ar.font.color.rgb = RGBColor(191, 97, 14)
ap.add_run('"Sau khi kiểm soát các biến giới tính, tuổi, áp lực học tập, thu nhập gia đình, HS ngủ ít có odds lo âu cao gấp AOR lần HS ngủ đủ" — với "AOR" là con số sau điều chỉnh (thường NHỎ HƠN OR thô).')
d.add_paragraph()

# =========================================================
# PHẦN 3 — CÁCH TÍNH AOR
# =========================================================
add_heading('3. Cách tính AOR qua hồi quy logistic đa biến', level=1)

add_para('AOR KHÔNG tính trực tiếp từ bảng 2×2. Phải dùng HỒI QUY LOGISTIC ĐA BIẾN (multivariable logistic regression).', bold=True)
d.add_paragraph()

add_para('3.1. Mô hình hồi quy logistic', bold=True, size=12)
add_para('Công thức:', bold=True)
add_para('log(odds) = β₀ + β₁·X₁ + β₂·X₂ + β₃·X₃ + ... + βₖ·Xₖ', bold=True, size=13, color=(192, 0, 0))
d.add_paragraph()

add_para('Trong đó:')
for it in [
    'log(odds) = log-odds của outcome (lo âu) — còn gọi là logit',
    'β₀ = hệ số chặn (intercept)',
    'β₁ = hệ số cho biến chính X₁ (ví dụ ngủ <5h)',
    'β₂, β₃... = hệ số cho các biến covariate (tuổi, giới, áp lực học...)',
    'Mô hình được fit bằng phương pháp MAXIMUM LIKELIHOOD — phần mềm (R, SPSS, Stata) tự giải',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

add_para('3.2. Từ β sang AOR', bold=True, size=12)
add_para('AOR = e^β (exponential của hệ số beta)', bold=True, size=13, color=(192, 0, 0))
d.add_paragraph()

add_para('Phần mềm thống kê sẽ output cho thầy:')
for it in [
    'β (coefficient) — ví dụ β₁ = 2,62',
    'SE (standard error) — ví dụ SE₁ = 0,15',
    'AOR = exp(β) — ví dụ exp(2,62) = 13,71',
    'KTC 95% của AOR = exp(β ± 1,96 × SE) — ví dụ exp(2,62 ± 1,96 × 0,15) = [10,2; 18,4]',
    'p-value từ Wald test = β / SE rồi so với phân phối chuẩn',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs:
        r.font.size = Pt(11)

d.add_paragraph()

# =========================================================
# PHẦN 4 — VÍ DỤ TÍNH TỪNG BƯỚC
# =========================================================
add_heading('4. Ví dụ tính AOR từng bước — Zhu 2025 (QT05)', level=1)

add_para('Zhu và cộng sự (2025) nghiên cứu lo âu HS Trung Quốc Suzhou với DASS-21. Outcome: lo âu (có/không). Biến chính: giấc ngủ <5h. Covariates: tuổi, giới, áp lực học, gia đình đơn thân.', bold=False)
d.add_paragraph()

add_para('BƯỚC 1 — Thu thập dữ liệu', bold=True)
add_para('Với mỗi HS, ghi nhận: lo âu (0/1), ngủ <5h (0/1), tuổi, giới, áp lực (điểm), gia đình đơn thân (0/1).')
d.add_paragraph()

add_para('BƯỚC 2 — Chạy logistic regression đa biến', bold=True)
add_para('Ví dụ lệnh trong R:')
code_tbl = d.add_table(rows=1, cols=1)
code_tbl.style = 'Table Grid'
cd = code_tbl.rows[0].cells[0]
shade_cell(cd, 'F2F2F2')
cp = cd.paragraphs[0]
cr = cp.add_run('model <- glm(anxiety ~ sleep_less5h + age + gender + pressure + single_parent,\n               data = df, family = binomial(link = "logit"))\nsummary(model)\nexp(coef(model))  # ra AOR\nexp(confint(model))  # ra KTC 95%')
cr.font.name = 'Consolas'
cr.font.size = Pt(10)
d.add_paragraph()

add_para('BƯỚC 3 — Đọc output', bold=True)
out_tbl = d.add_table(rows=6, cols=5)
out_tbl.style = 'Table Grid'
hdr = out_tbl.rows[0]
for i, h in enumerate(['Biến', 'β (coef)', 'SE', 'AOR = exp(β)', 'KTC 95%']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

out_rows = [
    ('Ngủ < 5h', '2,62', '0,15', '13,71', '[10,2; 18,4]'),
    ('Tuổi (mỗi năm)', '0,08', '0,04', '1,08', '[1,00; 1,17]'),
    ('Nữ (vs nam)', '0,35', '0,12', '1,42', '[1,12; 1,80]'),
    ('Áp lực học (mỗi điểm)', '0,12', '0,03', '1,13', '[1,06; 1,20]'),
    ('Gia đình đơn thân', '0,36', '0,10', '1,43', '[1,17; 1,76]'),
]
for i, row_data in enumerate(out_rows):
    row = out_tbl.rows[i+1]
    for j, v in enumerate(row_data):
        row.cells[j].text = v
        if j == 0:
            for p in row.cells[j].paragraphs:
                for r in p.runs:
                    r.bold = True
d.add_paragraph()

add_para('(Con số minh hoạ, AOR=13,71 và 1,43 là từ Zhu 2025 gốc; các cột khác là ví dụ.)', color=(128, 128, 128))
d.add_paragraph()

add_para('BƯỚC 4 — Diễn giải', bold=True)
inter_tbl = d.add_table(rows=4, cols=2)
inter_tbl.style = 'Table Grid'
inters = [
    ('AOR ngủ<5h = 13,71',
     'SAU KHI điều chỉnh tuổi, giới, áp lực, gia đình, HS ngủ <5h vẫn có odds lo âu cao gấp 13,71 lần HS ngủ đủ. Cho thấy giấc ngủ là yếu tố ĐỘC LẬP, không phải do biến khác gây nhiễu.'),
    ('AOR tuổi = 1,08',
     'Mỗi năm tuổi tăng thêm → odds lo âu tăng 8% (giữ cố định các biến khác).'),
    ('AOR nữ = 1,42',
     'Nữ có odds lo âu cao gấp 1,42 lần nam (đã kiểm soát tuổi, ngủ, áp lực...).'),
    ('AOR gia đình đơn thân = 1,43',
     'HS gia đình đơn thân có odds lo âu cao gấp 1,43 lần gia đình đủ cha mẹ.'),
]
for i, (k, v) in enumerate(inters):
    row = inter_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'E2EFDA')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(54, 95, 44)
    row.cells[1].text = v
d.add_paragraph()

# =========================================================
# PHẦN 5 — NGƯỠNG ĐỌC AOR
# =========================================================
add_heading('5. Ngưỡng đọc AOR', level=1)

th_tbl = d.add_table(rows=5, cols=2)
th_tbl.style = 'Table Grid'
hdr = th_tbl.rows[0]
for i, h in enumerate(['AOR', 'Ý nghĩa']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

ths = [
    ('AOR = 1', 'KHÔNG có mối liên hệ. Biến X không ảnh hưởng outcome sau khi kiểm soát covariates.'),
    ('AOR > 1', 'Biến X LÀM TĂNG odds outcome. Ví dụ AOR = 2 → gấp đôi odds. AOR = 13,71 → cực mạnh.'),
    ('AOR < 1', 'Biến X LÀM GIẢM odds outcome — yếu tố BẢO VỆ. Ví dụ AOR = 0,75 (Islam 2025: "cha mẹ quan tâm") → giảm 25 % odds lo âu.'),
    ('KTC 95 % chứa 1', 'KHÔNG có ý nghĩa thống kê. Nếu KTC = [0,8; 1,3] → chưa loại trừ AOR = 1 → không kết luận được.'),
]
for i, (k, v) in enumerate(ths):
    row = th_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FFE699')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.color.rgb = RGBColor(191, 97, 14)
    row.cells[1].text = v
d.add_paragraph()

# =========================================================
# PHẦN 6 — VÍ DỤ TỪ DỰ ÁN
# =========================================================
add_heading('6. Ví dụ AOR từ các bài trong thư viện dự án', level=1)

proj_tbl = d.add_table(rows=7, cols=3)
proj_tbl.style = 'Table Grid'
hdr = proj_tbl.rows[0]
for i, h in enumerate(['Bài', 'AOR (biến → outcome)', 'Diễn giải']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

projs = [
    ('QT05 Zhu 2025 TQ',
     'Ngủ <5h → lo âu: AOR = 13,71',
     'Yếu tố nguy cơ MẠNH NHẤT của lo âu tại Trung Quốc. Kiểm soát tuổi, giới, áp lực. Cực mạnh.'),
    ('QT05 Zhu 2025 TQ',
     'Gia đình đơn thân → lo âu: AOR = 1,434',
     'Yếu tố nguy cơ NHẸ. 43% odds cao hơn gia đình đủ cha mẹ.'),
    ('QT23 JAACAP 2024 Mỹ',
     'Lo âu trẻ em 2021 vs 2013: AOR = 2,17',
     'Lo âu TĂNG GẤP ĐÔI sau 8 năm, đã kiểm soát tuổi, giới, sắc tộc, bảo hiểm.'),
    ('QT31 Islam 2025 59 nước',
     'Bất an thực phẩm → lo âu: AOR = 2,22',
     'HS không đủ ăn có odds lo âu gấp 2,22 lần.'),
    ('QT31 Islam 2025 59 nước',
     'Ý định tự tử → lo âu: AOR = 2,84',
     'Yếu tố co-occurring mạnh nhất trong LMIC.'),
    ('QT31 Islam 2025 59 nước',
     'Cha mẹ quan tâm → lo âu: AOR = 0,75',
     'AOR < 1 → YẾU TỐ BẢO VỆ. Giảm 25 % odds lo âu.'),
]
for i, (k, v, n) in enumerate(projs):
    row = proj_tbl.rows[i+1]
    row.cells[0].text = k
    row.cells[1].text = v
    row.cells[2].text = n
    for p in row.cells[0].paragraphs:
        for r in p.runs:
            r.bold = True
d.add_paragraph()

# =========================================================
# PHẦN 7 — OR vs AOR vs RR
# =========================================================
add_heading('7. Phân biệt OR thô / AOR / RR', level=1)

cmp_tbl = d.add_table(rows=4, cols=3)
cmp_tbl.style = 'Table Grid'
hdr = cmp_tbl.rows[0]
for i, h in enumerate(['Chỉ số', 'Cách tính', 'Khi nào dùng']):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

cmps = [
    ('OR thô (crude)', 'Từ bảng 2×2 (a·d)/(b·c)', 'Cross-tab đơn biến. Không kiểm soát confounder.'),
    ('AOR (adjusted)', 'Từ logistic regression ĐA BIẾN. AOR = exp(β).', 'Đã kiểm soát các covariate. CHUẨN trong NC dịch tễ hiện đại.'),
    ('RR (Risk Ratio)', 'P(outcome | exposed) / P(outcome | unexposed)', 'Trong cohort study. RR gần OR khi outcome HIẾM (<10%). Khi outcome phổ biến → OR phóng đại RR.'),
]
for i, (k, v, n) in enumerate(cmps):
    row = cmp_tbl.rows[i+1]
    cell0 = row.cells[0]
    shade_cell(cell0, 'D9E1F2')
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(k)
    r0.bold = True
    row.cells[1].text = v
    row.cells[2].text = n
d.add_paragraph()

# Lỗi phổ biến
err_tbl = d.add_table(rows=1, cols=1)
err_tbl.style = 'Table Grid'
ec = err_tbl.rows[0].cells[0]
shade_cell(ec, 'FCE4D6')
ep = ec.paragraphs[0]
er = ep.add_run('⚠ LỖI PHỔ BIẾN: ')
er.bold = True
er.font.color.rgb = RGBColor(192, 0, 0)
ep.add_run('"AOR = 2 nghĩa là nguy cơ gấp đôi" → SAI KỸ THUẬT. AOR đo ODDS, không đo RISK. Khi outcome HIẾM (<10%), OR ≈ RR nên diễn giải "gấp đôi" tạm chấp nhận. Nhưng khi outcome PHỔ BIẾN (>20%, như lo âu HS 40-50%), AOR phóng đại RR. Ví dụ AOR=2 có thể tương đương RR chỉ 1,5. Nên nói chính xác: "AOR = 2 → odds cao gấp 2 lần", không phải "nguy cơ gấp 2".')
d.add_paragraph()

# =========================================================
# TÓM GỌN
# =========================================================
add_heading('8. Tóm gọn 3 ý cốt lõi', level=1)

core_tbl = d.add_table(rows=3, cols=1)
core_tbl.style = 'Table Grid'
cores = [
    ('① AOR = e^β từ logistic regression đa biến',
     'Phần mềm (R/SPSS/Stata) fit β cho mỗi biến độc lập bằng maximum likelihood. Lấy exp(β) được AOR. KTC 95% = exp(β ± 1,96·SE).',
     'E2EFDA'),
    ('② AOR kiểm soát confounder — tốt hơn OR thô',
     'Nếu chỉ báo cáo OR thô → rủi ro kết luận sai (ngủ ít "gây" lo âu trong khi thực tế áp lực học mới là nguyên nhân chính). AOR nói: "sau khi kiểm soát X, Y, Z, biến này vẫn có tác động RIÊNG bằng AOR".',
     'FFF2CC'),
    ('③ Đọc AOR với 3 con số',
     'AOR > 1: yếu tố nguy cơ. AOR < 1: yếu tố bảo vệ. KTC 95% chứa 1: không có ý nghĩa thống kê. Luôn đọc cả 3 (AOR, KTC, p-value) — không đọc riêng AOR.',
     'DEEBF7'),
]
for i, (k, v, clr) in enumerate(cores):
    cell = core_tbl.rows[i].cells[0]
    shade_cell(cell, clr)
    p = cell.paragraphs[0]
    r1 = p.add_run(k + '\n')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    r2 = p.add_run(v)
    r2.font.size = Pt(11)
d.add_paragraph()

# TLTK
add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Hosmer DW, Lemeshow S, Sturdivant RX. (2013). Applied Logistic Regression (3rd ed.). Wiley. [Giáo khoa chuẩn về logistic regression]',
    'Szumilas M. (2010). Explaining odds ratios. Journal of the Canadian Academy of Child and Adolescent Psychiatry, 19(3):227-229. [Bài ngắn gọn giải thích OR/AOR cho lâm sàng tâm thần trẻ em]',
    'Zhang J, Yu KF. (1998). What\'s the relative risk? A method of correcting the odds ratio in cohort studies of common outcomes. JAMA, 280(19):1690-1691. [Cảnh báo OR phóng đại RR khi outcome phổ biến]',
    'Zhu và cộng sự (2025). Anxiety and depression among high school students in China. BMC Public Health. [QT05 — nguồn AOR = 13,71 cho giấc ngủ]',
    'Mojtabai R, Olfson M. (2024). Trends in US children\'s mental disorders. JAACAP. [QT23 — nguồn AOR = 2,17]',
    'Islam và cộng sự (2025). Anxiety among adolescents in 59 LMICs. Journal of Affective Disorders. [QT31 — nguồn AOR = 2,22, 2,84, 0,75]',
    'STROBE Statement (2007). Strengthening the Reporting of Observational Studies in Epidemiology. [Chuẩn báo cáo AOR + KTC 95 %]',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 20/04/2026 | AOR=13,71 (Zhu 2025), 2,17 (JAACAP 2024), 2,22/2,84/0,75 (Islam 2025) lấy từ INSIGHT_05 và các tóm tắt QT05, QT23, QT31 trong thư viện. Công thức AOR = exp(β) là chuẩn trong logistic regression (Hosmer-Lemeshow 2013).')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/AOR_Adjusted_Odds_Ratio_cach_tinh.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
