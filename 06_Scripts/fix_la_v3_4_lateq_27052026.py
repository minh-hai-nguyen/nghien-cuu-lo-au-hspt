# -*- coding: utf-8 -*-
"""Fix LA v3_3 -> v3_4: remaining abbreviations LA tổng quát + LATQ.
27/05/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_3_FixSaikiaXu_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_4_FixAbbrev_27052026.docx')
RED = RGBColor(192, 0, 0)

# Careful order: longer first
REPLACEMENTS = [
    ('RLLA tổng quát', 'RLLA lan tỏa'),
    ('LA tổng quát', 'LA lan tỏa'),
    ('LATQ', 'LALT'),  # but NOT change RLLATQ (already fixed to RLLALT)
]


def apply_repl(p, replacements, applied, idx):
    full = p.text
    if not full:
        return False
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            i = full.find(old, cursor)
            if i == -1: break
            # Special check for "LATQ" - skip if preceded by "RL" (already RLLATQ -> RLLALT done)
            if old == 'LATQ' and i >= 2 and full[i-2:i] == 'RL':
                cursor = i + 1
                continue
            matches.append((i, i + len(old), old, new))
            cursor = i + len(old)
    if not matches:
        return False
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    safe = []
    last = -1
    for m in matches:
        if m[0] >= last:
            safe.append(m)
            last = m[1]
    segments = []
    cur = 0
    for s, e, o, n in safe:
        if cur < s:
            segments.append((full[cur:s], False))
        segments.append((n, True))
        applied.append(f"para {idx}: '{o}' -> '{n}'")
        cur = e
    if cur < len(full):
        segments.append((full[cur:], False))
    for r in p.runs:
        r.text = ''
    for text, is_red in segments:
        if not text: continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


print(f"=== Fix LA v3_3 -> v3_4 ===")
d = Document(SRC)
applied = []
for i, p in enumerate(d.paragraphs):
    apply_repl(p, REPLACEMENTS, applied, i)
for ti, tb in enumerate(d.tables):
    for ri, row in enumerate(tb.rows):
        for ci, cell in enumerate(row.cells):
            for cpi, p in enumerate(cell.paragraphs):
                apply_repl(p, REPLACEMENTS, applied, f'T{ti+1}r{ri}c{ci}p{cpi}')
d.save(OUT)
print(f"Saved: {OUT}")
print(f"Total: {len(applied)} fixes")
for a in applied:
    print(f"  ✓ {a}")
