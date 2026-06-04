# -*- coding: utf-8 -*-
"""Tạo file Word: 6 bài tham khảo đã xác minh, chuẩn bị viết bài Q1."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/Bài tham khảo đã xác minh - chuẩn bị bài Q1.docx"
BLUE = RGBColor(0x00, 0x00, 0xCC)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(13)


def H(t, sz=13):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(sz)
    return p


def body(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def field(label, t, blue=False):
    p = doc.add_paragraph()
    r0 = p.add_run(label + " ")
    r0.bold = True
    r = p.add_run(t)
    if blue:
        r.font.color.rgb = BLUE
        r.bold = True
    return p


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("SÁU BÀI THAM KHẢO ĐÃ XÁC MINH — CHUẨN BỊ VIẾT BÀI Q1")
tr.bold = True
tr.font.size = Pt(14)

body("Sáu bài dưới đây (3 bài Việt Nam + 3 bài quốc tế) được thu thập để chuẩn bị "
     "viết bài Q1 nối tiếp hướng nghiên cứu đường dẫn yếu tố nguy cơ/bảo vệ học "
     "đường đến lo âu vị thành niên. Toàn bộ thông tin thư mục, abstract và số liệu "
     "đã được mở và xác minh trực tiếp từ trang nguồn thật (PubMed, PubMed Central, "
     "trang nhà xuất bản) — không suy đoán, không bịa. Mỗi bài kèm đường dẫn nguồn "
     "để thầy/cô đối chiếu lần cuối trước khi đưa vào Danh mục tài liệu tham khảo.")
doc.add_paragraph()

H("A. BA BÀI VIỆT NAM")
doc.add_paragraph()

# VN1
H("A1. Tran và cộng sự (2024) — nghiên cứu DỌC 3 năm (ưu tiên cao nhất)")
field("Trích dẫn:", "Tran, T. V., Nguyen, H. T. L., Tran, X. M. T., Tashiro, Y., "
      "Seino, K., Vo, T. V., & Nakamura, K. (2024). Academic stress among students "
      "in Vietnam: A three-year longitudinal study on the impact of family, "
      "lifestyle, and academic factors. Journal of Rural Medicine, 19(4), 279–290.")
field("DOI / Nguồn:", "10.2185/jrm.2024-012 — PMC11442093 "
      "(https://pmc.ncbi.nlm.nih.gov/articles/PMC11442093/); J-STAGE.")
field("Thiết kế & mẫu:", "Nghiên cứu dọc 3 năm (2018–2021), 2 đợt đo; 4 trường "
      "THCS tại TP. Huế; 611 học sinh tuyển ban đầu, 341 hoàn thành cả hai đợt; "
      "lớp 6 (≈11 tuổi) → lớp 9. Công cụ: thang ESSA (16 mục, 5 tiểu thang).")
field("Phát hiện chính:", "Điểm ESSA trung bình tăng 46,4 ± 7,6 → 53,5 ± 10,8. "
      "Liên hệ với stress cao hơn: giới nữ, nhiều anh chị em (β = 2,24), học vấn "
      "cha cao hơn (β = 3,20), học thêm (β = 4,73); điểm học tập cao là yếu tố bảo "
      "vệ (β = −1,79). Yếu tố lối sống ít liên hệ.")
field("Truy cập mở:", "CÓ — toàn văn miễn phí (CC BY-NC-ND 4.0).")
field("→ Giá trị cho bài Q1:",
      "Một trong rất ít nghiên cứu DỌC về stress học đường ở HS THCS Việt Nam — "
      "nền tảng biện minh cho thiết kế đo lặp và lập luận diễn tiến theo thời gian "
      "của bài Q1.", blue=True)
doc.add_paragraph()

# VN2
H("A2. Pham và cộng sự (2024) — lo âu đặc thù, mẫu lớn")
field("Trích dẫn:", "Pham, T. T. H., Do, T. T., Nguyen, T. L., & Ngo, A. V. (2024). "
      "Anxiety symptoms and coping strategies among high school students in Vietnam "
      "after COVID-19 pandemic: A mixed-method evaluation. Frontiers in Public "
      "Health, 12, Article 1232856.")
field("DOI / Nguồn:", "10.3389/fpubh.2024.1232856 — PMC10904568 "
      "(https://pmc.ncbi.nlm.nih.gov/articles/PMC10904568/).")
field("Thiết kế & mẫu:", "Hỗn hợp (cắt ngang định lượng + phỏng vấn sâu); 3.910 "
      "học sinh THPT 14–17 tuổi (lớp 10–12), Hà Nội. Công cụ lo âu: GAD-7.")
field("Phát hiện chính:", "Tỷ lệ có triệu chứng lo âu 40,6% (nhẹ 23,9%; vừa 10,9%; "
      "nặng 5,8%). Nguồn gây lo âu (định tính): kết quả học tập, tương tác xã hội, "
      "định kiến, kỳ vọng gia đình.")
field("Truy cập mở:", "CÓ — toàn văn miễn phí (CC BY).")
field("→ Giá trị cho bài Q1:",
      "Bài đặc thù về LO ÂU (không phải trầm cảm), mẫu lớn, dùng GAD-7 — cung cấp "
      "tỷ lệ nền và nguồn nguy cơ học đường/gia đình cho phần Đặt vấn đề.", blue=True)
doc.add_paragraph()

# VN3
H("A3. Nguyen và cộng sự (2024) — khung yếu tố nguy cơ/bảo vệ ba tầng")
field("Trích dẫn:", "Nguyen, H. T. T., Tran, B. X., Luu, H. N., Boyer, L., Fond, "
      "G., Auquier, P., Latkin, C. A., Nguyen, T. T., Zhang, M. W. B., Ho, R. C. M., "
      "& Ho, C. S. H. (2024). Prevalence of depressive symptoms among urban school "
      "adolescents in Vietnam: The role of youth, family, and school relationships. "
      "Journal of Epidemiology and Population Health, 72(5), Article 202758.")
field("DOI / Nguồn:", "10.1016/j.jeph.2024.202758 — PubMed 39098167 "
      "(https://pubmed.ncbi.nlm.nih.gov/39098167/).")
field("Thiết kế & mẫu:", "Cắt ngang (2022); 507 học sinh THPT 15–17 tuổi, Hà Nội. "
      "Công cụ: thang trầm cảm RADS-2; hồi quy Tobit đa biến.")
field("Phát hiện chính:", "Yếu tố nguy cơ: quan hệ cha mẹ kém, không có người tâm "
      "sự, xung đột gia đình, áp lực thi cử. Yếu tố bảo vệ: hài lòng tình bạn ở "
      "trường, được thầy cô/bạn bè hỗ trợ, hoạt động ngoại khóa.")
field("Truy cập mở:", "Cần xác nhận trên trang ScienceDirect (PubMed không nêu PMC ID).")
field("→ Giá trị cho bài Q1:",
      "Phân tách rõ yếu tố nguy cơ vs bảo vệ ở ba tầng cá nhân–gia đình–nhà trường "
      "— trùng khít khung đường dẫn của bài Q1.", blue=True)
doc.add_paragraph()

H("B. BA BÀI QUỐC TẾ")
doc.add_paragraph()

# QT1
H("B1. Wang và cộng sự (2025) — phân tích mạng cắt-trễ DỌC")
field("Trích dẫn:", "Wang, J., Wang, T., & Cheng, Y. (2025). Longitudinal dynamics "
      "of anxiety and depression in adolescents with a history of childhood "
      "maltreatment: A cross-lagged panel network analysis. Stress and Health, "
      "41(4), Article e70076.")
field("DOI / Nguồn:", "10.1002/smi.70076 — PubMed 40643566 "
      "(https://pubmed.ncbi.nlm.nih.gov/40643566/).")
field("Thiết kế & mẫu:", "Nghiên cứu dọc 3 đợt (nền, 6 tháng, 12 tháng); 999 vị "
      "thành niên có tiền sử ngược đãi thời thơ ấu. Mô hình mạng bảng cắt-trễ.")
field("Phát hiện chính:", "Vấn đề quan hệ liên cá nhân và lo âu xã hội là 'triệu "
      "chứng cầu nối' nhất quán giữa cụm lo âu và trầm cảm qua các đợt; lo âu xã "
      "hội, triệu chứng cơ thể và né tránh tổn hại là yếu tố trung tâm.")
field("→ Giá trị cho bài Q1:",
      "Mẫu hình phương pháp DỌC + phân tích mạng — gợi ý nâng cấp từ SEM cắt ngang "
      "lên thiết kế dọc/cross-lagged cho bài Q1.", blue=True)
doc.add_paragraph()

# QT2
H("B2. Xiong và cộng sự (2024) — mô hình thác phát triển RI-CLPM")
field("Trích dẫn:", "Xiong, J., Fang, X., Wang, J., Xie, W., Liu, M., & Niu, G. "
      "(2024). Family cumulative risk, life satisfaction, and anxiety and "
      "depression in adolescents: A developmental cascades model. Journal of "
      "Adolescence, 96(7), 1445–1457.")
field("DOI / Nguồn:", "10.1002/jad.12354 — PubMed 38783637 "
      "(https://pubmed.ncbi.nlm.nih.gov/38783637/).")
field("Thiết kế & mẫu:", "Nghiên cứu dọc 3 đợt (10/2018, 4/2019, 11/2019); 707 học "
      "sinh THPT lớp 10–11, Vũ Hán, Trung Quốc. Mô hình bảng cắt-trễ chặn ngẫu "
      "nhiên (RI-CLPM).")
field("Phát hiện chính:", "Hiệu ứng cắt-trễ hai chiều: sự hài lòng cuộc sống và "
      "lo âu–trầm cảm dự báo nguy cơ tích lũy gia đình theo thời gian — ủng hộ góc "
      "nhìn 'đứa trẻ có ảnh hưởng' trong thác phát triển.")
field("→ Giá trị cho bài Q1:",
      "Mẫu hình RI-CLPM tách hiệu ứng trong-cá-nhân vs giữa-cá-nhân — khuôn mẫu "
      "phân tích để bài Q1 kiểm định chiều nhân quả giữa yếu tố gia đình và lo âu.",
      blue=True)
doc.add_paragraph()

# QT3
H("B3. Pan và cộng sự (2024) — SEM trung gian chuỗi, lo âu xã hội")
field("Trích dẫn:", "Pan, W., Li, B., Long, Y., & Cao, C. (2024). The relationship "
      "between perceived social support and social anxiety in Chongqing rural "
      "secondary school students: The chain mediating effect of core self-"
      "evaluation and shyness. BMC Psychology, 12(1), Article 708.")
field("DOI / Nguồn:", "10.1186/s40359-024-02229-z — PubMed 39614328 "
      "(https://pubmed.ncbi.nlm.nih.gov/39614328/).")
field("Thiết kế & mẫu:", "Cắt ngang; 626 học sinh trung học nông thôn, Trùng "
      "Khánh, Trung Quốc. Công cụ: PSSS, CSES, SS, SADS. Phân tích: mô hình cấu "
      "trúc tuyến tính (SEM), trung gian chuỗi.")
field("Phát hiện chính:", "Hỗ trợ xã hội tri giác và tự đánh giá cốt lõi là yếu tố "
      "BẢO VỆ khỏi lo âu xã hội; nhút nhát là yếu tố NGUY CƠ. Hỗ trợ xã hội tác "
      "động gián tiếp lên lo âu xã hội qua chuỗi tự đánh giá cốt lõi → nhút nhát; "
      "mô hình nhất quán giữa hai giới.")
field("→ Giá trị cho bài Q1:",
      "Khuôn mẫu SEM trung gian chuỗi trong bối cảnh tập thể Á Đông — đối chứng "
      "trực tiếp cho đường dẫn hỗ trợ xã hội → lo âu xã hội của bài Q1/Q2.5.",
      blue=True)
doc.add_paragraph()

H("GHI CHÚ XÁC MINH")
body("• Toàn bộ 6 bài đã được mở trực tiếp trang nguồn (PubMed/PMC/nhà xuất bản) "
     "và xác minh: tên đầy đủ tác giả, năm, tạp chí, tập/số/trang hoặc số bài, "
     "DOI, abstract, cỡ mẫu, thiết kế.")
body("• Hai mục cần đối chiếu nhanh khi đưa vào Danh mục TLTK chính thức: trạng "
     "thái truy cập mở của bài A3 (Nguyen 2024) và số trang chính xác — nên mở "
     "trang ScienceDirect xác nhận.")
body("• Các bài này MỚI là ứng viên tham khảo — chưa nhập vào thư viện canonical "
     "(chưa có mã VN/QT, chưa có bản dịch/tóm tắt). Nếu quyết định dùng, cần chạy "
     "quy trình nhập liệu đầy đủ.")

cp = doc.core_properties
cp.author = ""
cp.last_modified_by = ""
doc.save(OUT)
print("Đã lưu:", OUT)
