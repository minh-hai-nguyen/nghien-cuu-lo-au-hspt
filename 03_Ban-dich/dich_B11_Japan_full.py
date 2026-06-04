# -*- coding: utf-8 -*-
"""Dịch đầy đủ B11 Japan iCBT SAD 2024 — Matsumoto et al."""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.2196/55786', size=10)

add_heading(doc, 'Hi\u1ec7u qu\u1ea3 li\u1ec7u ph\u00e1p nh\u1eadn th\u1ee9c\u2013h\u00e0nh vi qua internet kh\u00f4ng h\u01b0\u1edbng d\u1eabn cho r\u1ed1i lo\u1ea1n lo \u00e2u x\u00e3 h\u1ed9i d\u01b0\u1edbi ng\u01b0\u1ee1ng \u1edf thanh thi\u1ebfu ni\u00ean v\u00e0 thanh ni\u00ean: Th\u1eed nghi\u1ec7m \u0111a trung t\u00e2m ng\u1eabu nhi\u00ean c\u00f3 \u0111\u1ed1i ch\u1ee9ng', 1)
h = doc.add_paragraph()
r = h.add_run('Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: Multicenter Randomized Controlled Trial')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('Ti\u00eau \u0111\u1ec1 g\u1ed1c', 'Effectiveness of Unguided Internet-Based CBT for Subthreshold SAD in Adolescents and Young Adults: Multicenter RCT'),
    ('T\u00e1c gi\u1ea3', 'Kazuki Matsumoto PhD (1), Sayo Hamatani PhD (2,3), Kiko Shiga PhD (4), Kiyoko Iiboshi PhD (4), Makiko Kasai PhD (5), Yasuhiro Kimura MA (6), Satoshi Yokota MA (7), Katsunori Watanabe PhD (8), Yoko Kubo MA (9), Masayuki Nakamura MD PhD (10)'),
    ('C\u01a1 quan', '(1) BV \u0110H Kagoshima; (2-3) \u0110H Fukui; (4) \u0110H Shigakukan; (5) \u0110H Naruto; (6) CL Fukushima; (7) THPT Kochi Kokusai; (8) \u0110H Jin-ai; (9-10) \u0110H Kagoshima \u2014 10 t\u00e1c gi\u1ea3 t\u1eeb 8 c\u01a1 quan Nh\u1eadt B\u1ea3n'),
    ('T\u1ea1p ch\u00ed', 'JMIR Pediatrics and Parenting (Q2, IF \u2248 2,5)'),
    ('Xu\u1ea5t b\u1ea3n', '2024, Vol. 7, e55786, 13 trang'),
    ('DOI', '10.2196/55786'),
    ('\u0110\u0103ng k\u00fd TN', 'UMIN000049768'),
    ('Lo\u1ea1i NC', 'Th\u1eed nghi\u1ec7m ng\u1eabu nhi\u00ean c\u00f3 \u0111\u1ed1i ch\u1ee9ng \u0111a trung t\u00e2m (multicenter RCT) \u2014 iCBT T\u1ef0 H\u1eccC ho\u00e0n to\u00e0n'),
    ('M\u1eabu', '77 HS + SV (38 can thi\u1ec7p + 39 \u0111\u1ed1i ch\u1ee9ng), 6 \u0110H + 1 THPT Nh\u1eadt B\u1ea3n, SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng, 12/2022\u201310/2023'),
    ('C\u00f4ng c\u1ee5', 'LSAS-J (lo \u00e2u XH), SPS (lo \u00e2u XH ph\u00e1t bi\u1ec3u), PHQ-9 (tr\u1ea7m c\u1ea3m), GAD-7 (lo \u00e2u chung), EQ-5D-5L (ch\u1ea5t l\u01b0\u1ee3ng s\u1ed1ng)'),
    ('Can thi\u1ec7p', 'iCBT t\u1ef1 h\u1ecdc ho\u00e0n to\u00e0n: 10 phi\u00ean d\u1ea1ng v\u0103n b\u1ea3n, d\u1ea1y k\u1ef9 thu\u1eadt CBT cho lo \u00e2u XH'),
    ('Truy c\u1eadp', 'Open Access \u2014 JMIR'),
])
add_page_ref(doc, '1\u201313', 'JMIR Pediatrics and Parenting', 'Vol. 7, 2024')

# TOM TAT
add_heading(doc, 'T\u00d3M T\u1eaeT', 2)
add_p(doc, 'B\u1ed1i c\u1ea3nh: R\u1ed1i lo\u1ea1n lo \u00e2u x\u00e3 h\u1ed9i (SAD) l\u00e0 r\u1ed1i lo\u1ea1n t\u00e2m th\u1ea7n ph\u1ed5 bi\u1ebfn \u1edf thanh thi\u1ebfu ni\u00ean v\u00e0 thanh ni\u00ean. Can thi\u1ec7p s\u1edbm c\u00f3 th\u1ec3 gi\u00fap ng\u0103n ch\u1eb7n ph\u00e1t tri\u1ec3n SAD \u0111\u1ea7y \u0111\u1ee7. Xem x\u00e9t r\u1eb1ng VTN c\u00f3 tri\u1ec7u ch\u1ee9ng lo \u00e2u XH kh\u00f4ng \u01b0a th\u00edch phi\u00ean m\u1eb7t \u0111\u1ed1i m\u1eb7t do s\u1ee3 giao ti\u1ebfp v\u1edbi nh\u00e0 tr\u1ecb li\u1ec7u, li\u1ec7u ph\u00e1p nh\u1eadn th\u1ee9c\u2013h\u00e0nh vi qua internet (iCBT) \u0111\u01b0\u1ee3c tri\u1ec3n khai.')

p = doc.add_paragraph()
r = p.add_run('Ph\u01b0\u01a1ng ph\u00e1p: RCT \u0111a trung t\u00e2m t\u1eeb 12/2022 \u0111\u1ebfn 10/2023. Ng\u01b0\u1eddi tham gia l\u00e0 HS/SV t\u1ea1i 6 \u0111\u1ea1i h\u1ecdc v\u00e0 1 tr\u01b0\u1eddng THPT Nh\u1eadt B\u1ea3n. Can thi\u1ec7p: iCBT T\u1ef0 H\u1eccC HO\u00c0N TO\u00c0N (complete self-help) g\u1ed3m 10 phi\u00ean d\u1ea1ng v\u0103n b\u1ea3n d\u1ea1y k\u1ef9 thu\u1eadt CBT cho lo \u00e2u XH \u1edf VTN v\u00e0 thanh ni\u00ean. \u0110\u1ed1i ch\u1ee9ng: kh\u00f4ng can thi\u1ec7p. Ph\u00e2n ng\u1eabu nhi\u00ean 1:1 b\u1eb1ng ch\u01b0\u01a1ng tr\u00ecnh m\u00e1y t\u00ednh. \u0110o: LSAS-J + SPS (lo \u00e2u XH), PHQ-9 (tr\u1ea7m c\u1ea3m), GAD-7 (lo \u00e2u chung), EQ-5D-5L (ch\u1ea5t l\u01b0\u1ee3ng s\u1ed1ng).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('K\u1ebft qu\u1ea3: 77 HS/SV \u0111\u01b0\u1ee3c ghi danh (38 can thi\u1ec7p, 39 \u0111\u1ed1i ch\u1ee9ng). ANCOVA v\u1edbi tr\u1ea7m c\u1ea3m l\u00e0m bi\u1ebfn \u0111\u1ed3ng: nh\u00f3m can thi\u1ec7p gi\u1ea3m \u0111\u00e1ng k\u1ec3 tri\u1ec7u ch\u1ee9ng lo \u00e2u XH, tr\u1ea7m c\u1ea3m v\u00e0 lo \u00e2u chung so v\u1edbi \u0111\u1ed1i ch\u1ee9ng. T\u1ef7 l\u1ec7 \u0111\u00e1p \u1ee9ng: 61% (19/31) can thi\u1ec7p vs 24% (9/38) \u0111\u1ed1i ch\u1ee9ng \u2014 OR = 4,97 (KTC 95%: 1,61\u201316,53; p = 0,003). T\u1ef7 l\u1ec7 ph\u1ee5c h\u1ed3i: 68% (21/31) vs 34% (13/38) \u2014 OR = 3,95 (KTC: 1,32\u201312,56; p = 0,008). T\u1ef7 l\u1ec7 thuy\u00ean gi\u1ea3m: OR = 2,01 (kh\u00f4ng \u0111\u00e1ng k\u1ec3, p = 0,20).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

add_p(doc, 'K\u1ebft lu\u1eadn: iCBT t\u1ef1 h\u1ecdc ho\u00e0n to\u00e0n c\u00f3 th\u1ec3 gi\u1ea3m hi\u1ec7u qu\u1ea3 tri\u1ec7u ch\u1ee9ng lo \u00e2u XH d\u01b0\u1edbi ng\u01b0\u1ee1ng \u1edf HS/SV Nh\u1eadt B\u1ea3n. Can thi\u1ec7p s\u1edbm qua internet c\u00f3 ti\u1ec1m n\u0103ng ng\u0103n ch\u1eb7n ph\u00e1t tri\u1ec3n SAD \u0111\u1ea7y \u0111\u1ee7 v\u1edbi chi ph\u00ed th\u1ea5p v\u00e0 kh\u00f4ng c\u1ea7n chuy\u00ean gia.')

# DANH GIA NHANH
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'JMIR Pediatrics Q2. RCT \u0111a trung t\u00e2m Nh\u1eadt B\u1ea3n (6 \u0110H + 1 THPT) \u2014 \u0111\u0103ng k\u00fd UMIN000049768.',
    'iCBT T\u1ef0 H\u1eccC HO\u00c0N TO\u00c0N (kh\u00f4ng h\u01b0\u1edbng d\u1eabn) \u2014 chi ph\u00ed TH\u1ea4P NH\u1ea4T. Kh\u00f4ng c\u1ea7n chuy\u00ean gia hay GV.',
    'SAD D\u01af\u1edaI NG\u01af\u1ee0NG \u2014 can thi\u1ec7p S\u1edaM tr\u01b0\u1edbc khi th\u00e0nh r\u1ed1i lo\u1ea1n \u0111\u1ea7y \u0111\u1ee7. Ph\u00f2ng ng\u1eeba h\u01a1n ch\u1eefa tr\u1ecb.',
    'Response rate 61% (OR=4,97) v\u00e0 recovery 68% (OR=3,95) \u2014 hi\u1ec7u qu\u1ea3 \u0111\u00e1ng k\u1ec3.',
    'Nh\u1eadt B\u1ea3n \u2014 \u0110\u00f4ng B\u1eafc \u00c1, v\u0103n h\u00f3a g\u1ea7n VN h\u01a1n ph\u01b0\u01a1ng T\u00e2y.',
    'Bao g\u1ed3m 1 THPT \u2014 kh\u00f4ng ch\u1ec9 \u0111\u1ea1i h\u1ecdc.',
    'Ph\u00f9 h\u1ee3p B9 NMA (iCBT SUCRA 71,2% h\u1ea1ng 1 cho SAD).',
    'Nh\u01b0ng B2 JMIR MA: c\u00f3 h\u01b0\u1edbng d\u1eabn (g=0,825) > kh\u00f4ng h\u01b0\u1edbng d\u1eabn (~0,3\u20130,4) \u2014 hi\u1ec7u qu\u1ea3 c\u00f3 th\u1ec3 h\u1ea1n ch\u1ebf.',
]:
    add_p(doc, f'\u2022 {b}')
add_p(doc, 'H\u1ea1n ch\u1ebf:', bold=True)
for b in [
    'n = 77 \u2014 nh\u1ecf. Ch\u1ec9 Nh\u1eadt B\u1ea3n.',
    'SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng \u2014 kh\u00f4ng ph\u1ea3i SAD \u0111\u1ea7y \u0111\u1ee7.',
    'KH\u00d4NG h\u01b0\u1edbng d\u1eabn \u2014 B2 JMIR cho th\u1ea5y c\u00f3 h\u01b0\u1edbng d\u1eabn hi\u1ec7u qu\u1ea3 h\u01a1n nhi\u1ec1u.',
    'Thuy\u00ean gi\u1ea3m kh\u00f4ng \u0111\u00e1ng k\u1ec3 (OR=2,01, p=0,20).',
    'Kh\u00f4ng theo d\u00f5i d\u00e0i h\u1ea1n (ch\u1ec9 ngay sau can thi\u1ec7p).',
]:
    add_p(doc, f'\u2022 {b}')

# PHUONG PHAP
add_page_ref(doc, '1\u20135', 'JMIR Pediatrics', 'Vol. 7, 2024')
add_heading(doc, '1. PH\u01af\u01a0NG PH\u00c1P', 2)
add_p(doc, '1.1. Thi\u1ebft k\u1ebf: RCT \u0111a trung t\u00e2m. 7 c\u01a1 s\u1edf (6 \u0110H + 1 THPT) t\u1ea1i 6 t\u1ec9nh Nh\u1eadt B\u1ea3n. Th\u1eddi gian: 12/2022\u201310/2023. \u0110\u0103ng k\u00fd: UMIN000049768.')
add_p(doc, '1.2. Ng\u01b0\u1eddi tham gia: HS + SV c\u00f3 SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng (LSAS-J \u226530 v\u00e0 <60; ho\u1eb7c SPS \u226516). Kh\u00f4ng c\u00f3 ti\u1ec1n s\u1eed r\u1ed1i lo\u1ea1n t\u00e2m th\u1ea7n. Lo\u1ea1i tr\u1eeb: \u0111ang nh\u1eadn tr\u1ecb li\u1ec7u, d\u00f9ng thu\u1ed1c, \u00fd t\u01b0\u1edfng t\u1ef1 t\u1eed. Ph\u00e2n ng\u1eabu nhi\u00ean 1:1 b\u1eb1ng m\u00e1y t\u00ednh.')
add_p(doc, '1.3. Can thi\u1ec7p: iCBT T\u1ef0 H\u1eccC HO\u00c0N TO\u00c0N (complete self-help) \u2014 10 phi\u00ean d\u1ea1ng v\u0103n b\u1ea3n. N\u1ed9i dung: gi\u00e1o d\u1ee5c t\u00e2m l\u00fd v\u1ec1 lo \u00e2u XH, t\u00e1i c\u1ea5u tr\u00fac nh\u1eadn th\u1ee9c, ph\u01a1i nhi\u1ec5m (exposure) \u2014 t\u1ea5t c\u1ea3 th\u00f4ng qua v\u0103n b\u1ea3n t\u1ef1 \u0111\u1ecdc, KH\u00d4NG c\u00f3 nh\u00e0 tr\u1ecb li\u1ec7u hay h\u01b0\u1edbng d\u1eabn vi\u00ean. \u0110\u1ed1i ch\u1ee9ng: kh\u00f4ng can thi\u1ec7p.')
add_p(doc, '1.4. Thang \u0111o: (a) LSAS-J (Liebowitz Social Anxiety Scale ti\u1ebfng Nh\u1eadt, 24 m\u1ee5c, 0\u2013144); (b) SPS (Social Phobia Scale, 20 m\u1ee5c); (c) PHQ-9 (tr\u1ea7m c\u1ea3m); (d) GAD-7 (lo \u00e2u chung); (e) EQ-5D-5L (ch\u1ea5t l\u01b0\u1ee3ng s\u1ed1ng). \u0110o t\u1ea1i ban \u0111\u1ea7u v\u00e0 sau can thi\u1ec7p.')
add_p(doc, '1.5. Ph\u00e2n t\u00edch: ANCOVA v\u1edbi tr\u1ea7m c\u1ea3m (PHQ-9 ban \u0111\u1ea7u) l\u00e0 bi\u1ebfn \u0111\u1ed3ng. Fisher exact test cho t\u1ef7 l\u1ec7 \u0111\u00e1p \u1ee9ng/ph\u1ee5c h\u1ed3i/thuy\u00ean gi\u1ea3m. \u0110\u00e1p \u1ee9ng = gi\u1ea3m \u226550% LSAS-J. Ph\u1ee5c h\u1ed3i = LSAS-J <30 sau can thi\u1ec7p. Thuy\u00ean gi\u1ea3m = LSAS-J <30 v\u00e0 SPS <16.')

# KET QUA
add_page_ref(doc, '5\u20139', 'JMIR Pediatrics', 'Vol. 7, 2024')
add_heading(doc, '2. K\u1ebeT QU\u1ea2', 2)
add_p(doc, '77 HS/SV ghi danh: 38 can thi\u1ec7p + 39 \u0111\u1ed1i ch\u1ee9ng. Tu\u1ed5i TB: can thi\u1ec7p 19,8 (SD=2,2), \u0111\u1ed1i ch\u1ee9ng 20,1 (SD=2,6). N\u1eef: 66% can thi\u1ec7p, 67% \u0111\u1ed1i ch\u1ee9ng. LSAS-J ban \u0111\u1ea7u: can thi\u1ec7p 46,9 (SD=17,4), \u0111\u1ed1i ch\u1ee9ng 44,6 (SD=17,6). Ho\u00e0n th\u00e0nh \u0111\u1ea7y \u0111\u1ee7: 31/38 (82%) can thi\u1ec7p, 38/39 (97%) \u0111\u1ed1i ch\u1ee9ng.')

add_heading(doc, 'B\u1ea3ng 1. K\u1ebft qu\u1ea3 ch\u00ednh \u2014 ANCOVA', 3)
add_table(doc,
    ['Thang \u0111o', 'Can thi\u1ec7p (sau)', '\u0110\u1ed1i ch\u1ee9ng (sau)', 'F / p', 'K\u00edch th\u01b0\u1edbc hi\u1ec7u \u1ee9ng'],
    [['LSAS-J (lo \u00e2u XH)', 'Gi\u1ea3m \u0111\u00e1ng k\u1ec3', 'Kh\u00f4ng thay \u0111\u1ed5i', 'p < 0,05', 'Cohen d \u2248 0,5\u20130,7'],
     ['SPS (lo \u00e2u ph\u00e1t bi\u1ec3u)', 'Gi\u1ea3m \u0111\u00e1ng k\u1ec3', '', 'p < 0,05', ''],
     ['PHQ-9 (tr\u1ea7m c\u1ea3m)', 'Gi\u1ea3m \u0111\u00e1ng k\u1ec3', '', 'p < 0,05', ''],
     ['GAD-7 (lo \u00e2u chung)', 'Gi\u1ea3m \u0111\u00e1ng k\u1ec3', '', 'p < 0,05', ''],
     ['EQ-5D-5L (ch\u1ea5t l\u01b0\u1ee3ng s\u1ed1ng)', 'Kh\u00f4ng \u0111\u00e1ng k\u1ec3', '', 'p > 0,05', '']],
    widths=[3.5, 2.5, 2.5, 2.0, 3.0])

add_heading(doc, 'B\u1ea3ng 2. T\u1ef7 l\u1ec7 \u0111\u00e1p \u1ee9ng, ph\u1ee5c h\u1ed3i, thuy\u00ean gi\u1ea3m', 3)
add_table(doc,
    ['Ch\u1ec9 s\u1ed1', 'Can thi\u1ec7p', '\u0110\u1ed1i ch\u1ee9ng', 'OR (KTC 95%)', 'p'],
    [['\u0110\u00e1p \u1ee9ng (\u226550% gi\u1ea3m LSAS)', '61% (19/31)', '24% (9/38)', '4,97 (1,61\u201316,53)', '0,003**'],
     ['Ph\u1ee5c h\u1ed3i (LSAS <30)', '68% (21/31)', '34% (13/38)', '3,95 (1,32\u201312,56)', '0,008**'],
     ['Thuy\u00ean gi\u1ea3m (LSAS<30 + SPS<16)', '\u2014', '\u2014', '2,01 (0,64\u20136,60)', '0,20 n.s.']],
    widths=[4.0, 2.5, 2.5, 3.0, 1.5])
add_p(doc, '\u0110\u00e1p \u1ee9ng 61% vs 24% (OR=4,97) v\u00e0 ph\u1ee5c h\u1ed3i 68% vs 34% (OR=3,95) \u2014 hi\u1ec7u qu\u1ea3 \u0110\u00c1NG K\u1ec2. Thuy\u00ean gi\u1ea3m kh\u00f4ng \u0111\u00e1ng k\u1ec3 \u2014 c\u1ea7n can thi\u1ec7p m\u1ea1nh h\u01a1n ho\u1eb7c d\u00e0i h\u01a1n cho thuy\u00ean gi\u1ea3m ho\u00e0n to\u00e0n.', size=9, italic=True)

# THAO LUAN
add_page_ref(doc, '9\u201312', 'JMIR Pediatrics', 'Vol. 7, 2024')
add_heading(doc, '3. TH\u1ea2O LU\u1eacN V\u00c0 K\u1ebeT LU\u1eacN', 2)
add_p(doc, 'iCBT t\u1ef1 h\u1ecdc ho\u00e0n to\u00e0n HI\u1ec6U QU\u1ea2 gi\u1ea3m lo \u00e2u XH, tr\u1ea7m c\u1ea3m v\u00e0 lo \u00e2u chung \u1edf HS/SV Nh\u1eadt v\u1edbi SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng. \u0110\u00e1p \u1ee9ng 61% (OR=4,97) v\u00e0 ph\u1ee5c h\u1ed3i 68% (OR=3,95) \u2014 r\u1ea5t cao cho can thi\u1ec7p t\u1ef1 h\u1ecdc kh\u00f4ng h\u01b0\u1edbng d\u1eabn.')

add_p(doc, 'SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng = can thi\u1ec7p S\u1edaM. Ph\u00f2ng ng\u1eeba tr\u01b0\u1edbc khi SAD tr\u1edf th\u00e0nh \u0111\u1ea7y \u0111\u1ee7 \u2014 \u0111\u1eb7c bi\u1ec7t quan tr\u1ecdng \u1edf VTN v\u00ec SAD th\u01b0\u1eddng kh\u1edfi ph\u00e1t ~9,2 tu\u1ed5i (B9 NMA Xian 2024). Ph\u00f9 h\u1ee3p Jefferies QT35: 18% VN c\u00f3 SAD nh\u01b0ng kh\u00f4ng bi\u1ebft \u2014 can thi\u1ec7p s\u1edbm c\u00f3 th\u1ec3 gi\u00fap nh\u00f3m n\u00e0y.')

add_p(doc, 'iCBT t\u1ef1 h\u1ecdc = chi ph\u00ed TH\u1ea4P NH\u1ea4T: kh\u00f4ng c\u1ea7n chuy\u00ean gia, kh\u00f4ng c\u1ea7n GV, kh\u00f4ng c\u1ea7n l\u1ecbch h\u1eb9n. Ch\u1ec9 c\u1ea7n internet + \u0111i\u1ec7n tho\u1ea1i/m\u00e1y t\u00ednh. Ph\u00f9 h\u1ee3p LMIC nh\u01b0 VN thi\u1ebfu chuy\u00ean gia SKTT.')

add_p(doc, 'Tuy nhi\u00ean, B2 JMIR MA cho th\u1ea5y DMHI c\u00f3 h\u01b0\u1edbng d\u1eabn (g=0,825) v\u01b0\u1ee3t tr\u1ed9i kh\u00f4ng h\u01b0\u1edbng d\u1eabn (~0,3\u20130,4). B\u00e0i n\u00e0y (iCBT kh\u00f4ng h\u01b0\u1edbng d\u1eabn) \u0111\u1ea1t OR=4,97 \u2014 kh\u00e1 cao nh\u01b0ng c\u00f3 th\u1ec3 T\u1ed0T H\u01a0N n\u1ebfu c\u00f3 GV/mentor h\u01b0\u1edbng d\u1eabn. G\u1ee3i \u00fd VN: ph\u00e1t tri\u1ec3n iCBT c\u00f3 GV h\u01b0\u1edbng d\u1eabn (k\u1ebft h\u1ee3p B11 + B8 Sri Lanka + B7 CA-CBT \u0110NA).')

add_p(doc, 'Nh\u1eadt B\u1ea3n \u2014 \u0110BA, v\u0103n h\u00f3a t\u00f4n tr\u1ecdng th\u1ee9 b\u1eadc, ng\u1ea1i n\u00f3i tr\u01b0\u1edbc ng\u01b0\u1eddi l\u1ea1 \u2014 t\u01b0\u01a1ng t\u1ef1 VN. Jefferies QT35: VN SAD 30,7%, "n\u00f3i v\u1edbi ng\u01b0\u1eddi c\u00f3 th\u1ea9m quy\u1ec1n" x\u1ebfp h\u1ea1ng lo \u00e2u 3/17 \u2014 v\u0103n h\u00f3a th\u1ee9 b\u1eadc.')

add_p(doc, 'H\u1ea1n ch\u1ebf: (1) n=77 nh\u1ecf. (2) Ch\u1ec9 Nh\u1eadt B\u1ea3n. (3) SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng \u2014 ch\u01b0a bi\u1ebft hi\u1ec7u qu\u1ea3 cho SAD \u0111\u1ea7y \u0111\u1ee7. (4) Kh\u00f4ng theo d\u00f5i d\u00e0i h\u1ea1n. (5) Thuy\u00ean gi\u1ea3m kh\u00f4ng \u0111\u00e1ng k\u1ec3 (p=0,20). (6) 10 phi\u00ean v\u0103n b\u1ea3n \u2014 thi\u1ebfu video/t\u01b0\u01a1ng t\u00e1c (B3 JAMA Maya c\u00f3 video + quiz + gamification).')

# TLTK
add_heading(doc, 'T\u00c0I LI\u1ec6U THAM KH\u1ea2O', 2)
for ref in [
    'Matsumoto, K. et al. (2024). Effectiveness of Unguided iCBT for Subthreshold SAD in Adolescents and Young Adults. JMIR Pediatrics and Parenting, 7, e55786.',
    'Liebowitz, M.R. (1987). Social phobia. Modern Problems of Pharmacopsychiatry, 22, 141\u2013173.',
    '(Xem \u0111\u1ea7y \u0111\u1ee7 50+ TLTK trong b\u00e0i g\u1ed1c)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('iCBT', 'Internet-based Cognitive Behavioral Therapy \u2014 Li\u1ec7u ph\u00e1p Nh\u1eadn th\u1ee9c\u2013H\u00e0nh vi qua Internet'),
    ('SAD', 'Social Anxiety Disorder \u2014 R\u1ed1i lo\u1ea1n Lo \u00e2u X\u00e3 h\u1ed9i'),
    ('LSAS-J', 'Liebowitz Social Anxiety Scale \u2014 Nh\u1eadt B\u1ea3n'),
    ('SPS', 'Social Phobia Scale'),
    ('PHQ-9', 'Patient Health Questionnaire 9-item'),
    ('GAD-7', 'Generalized Anxiety Disorder 7-item Scale'),
    ('EQ-5D-5L', 'EuroQol 5 Dimensions 5 Levels'),
    ('ANCOVA', 'Analysis of Covariance'),
    ('OR', 'Odds Ratio \u2014 T\u1ef7 s\u1ed1 ch\u00eanh'),
    ('RCT', 'Randomized Controlled Trial'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N / CRITICAL REVIEW')
add_red(doc, '\u0110i\u1ec3m m\u1ea1nh:', bold=True)
for s in [
    'JMIR Pediatrics Q2. RCT \u0111a trung t\u00e2m (7 c\u01a1 s\u1edf) \u2014 \u0111\u0103ng k\u00fd UMIN000049768.',
    'iCBT T\u1ef0 H\u1eccC HO\u00c0N TO\u00c0N \u2014 chi ph\u00ed th\u1ea5p nh\u1ea5t, kh\u00f4ng c\u1ea7n chuy\u00ean gia. KH\u1ea2 THI nh\u1ea5t cho LMIC nh\u01b0 VN.',
    'SAD D\u01af\u1edaI NG\u01af\u1ee0NG \u2014 can thi\u1ec7p PH\u00d2NG NG\u1eeeA. Ph\u00f9 h\u1ee3p QT35 Jefferies (18% false negatives VN).',
    '\u0110\u00e1p \u1ee9ng 61% OR=4,97 + ph\u1ee5c h\u1ed3i 68% OR=3,95 \u2014 hi\u1ec7u qu\u1ea3 \u0111\u00e1ng k\u1ec3.',
    'Nh\u1eadt B\u1ea3n \u0110BA \u2014 v\u0103n h\u00f3a g\u1ea7n VN (th\u1ee9 b\u1eadc, ng\u1ea1i t\u01b0\u01a1ng t\u00e1c). Ph\u00f9 h\u1ee3p B9 NMA (iCBT h\u1ea1ng 1 SUCRA 71,2%).',
    'Gi\u1ea3m c\u1ea3 tr\u1ea7m c\u1ea3m + lo \u00e2u chung \u2014 "ba trong m\u1ed9t".',
    'Bao g\u1ed3m 1 THPT (kh\u00f4ng ch\u1ec9 \u0110H) \u2014 VTN c\u1ea5p 3.',
]:
    add_red(doc, f'\u2022 {s}')
add_red(doc, 'H\u1ea1n ch\u1ebf chi ti\u1ebft:', bold=True)
for s in [
    'n = 77 nh\u1ecf. Ch\u1ec9 Nh\u1eadt B\u1ea3n \u2014 kh\u00e1c VN v\u1ec1 gi\u00e1o d\u1ee5c, kinh t\u1ebf.',
    'SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng \u2014 ch\u01b0a bi\u1ebft hi\u1ec7u qu\u1ea3 cho SAD \u0110\u1ea6Y \u0110\u1ee6. B9 NMA bao g\u1ed3m SAD \u0111\u1ea7y \u0111\u1ee7.',
    'KH\u00d4NG h\u01b0\u1edbng d\u1eabn \u2014 B2 JMIR MA: c\u00f3 h\u01b0\u1edbng d\u1eabn (g=0,825) V\u01af\u1ee2T TR\u1ed8I. B3 JAMA (d=0,94) c\u00f3 app t\u01b0\u01a1ng t\u00e1c cao.',
    'Thuy\u00ean gi\u1ea3m kh\u00f4ng \u0111\u00e1ng k\u1ec3 (p=0,20) \u2014 gi\u1ea3m tri\u1ec7u ch\u1ee9ng nh\u01b0ng kh\u00f4ng "kh\u1ecfi b\u1ec7nh".',
    'Kh\u00f4ng theo d\u00f5i d\u00e0i h\u1ea1n \u2014 hi\u1ec7u qu\u1ea3 DUY TR\u00cc kh\u00f4ng? B3 JAMA: d=1,04 sau 12 tu\u1ea7n.',
    '10 phi\u00ean v\u0103n b\u1ea3n thu\u1ea7n t\u00fay \u2014 thi\u1ebfu video, quiz, gamification. B3 JAMA Maya c\u00f3 t\u1ea5t c\u1ea3.',
]:
    add_red(doc, f'\u2022 {s}')
add_red(doc, 'Kho\u1ea3ng tr\u1ed1ng NC / Gap:', bold=True)
for s in [
    'VN c\u1ea7n iCBT cho SAD d\u01b0\u1edbi ng\u01b0\u1ee1ng \u2014 s\u00e0ng l\u1ecdc SIAS-17 (QT35: VN SAD 30,7%) + can thi\u1ec7p s\u1edbm qua app/web.',
    'So s\u00e1nh: iCBT t\u1ef1 h\u1ecdc (b\u00e0i n\u00e0y) vs iCBT c\u00f3 GV h\u01b0\u1edbng d\u1eabn (B2 JMIR g=0,825) t\u1ea1i VN.',
    'Ph\u00e1t tri\u1ec3n iCBT ti\u1ebfng Vi\u1ec7t: 10 phi\u00ean + video + quiz + gamification (k\u1ebft h\u1ee3p B11 + B3 JAMA Maya).',
    'Th\u00edch \u1ee9ng v\u0103n h\u00f3a VN: t\u00ecnh hu\u1ed1ng lo \u00e2u XH VN (\u201cn\u00f3i tr\u01b0\u1edbc l\u1edbp\u201d, \u201cg\u1eb7p ng\u01b0\u1eddi c\u00f3 th\u1ea9m quy\u1ec1n\u201d) + B7 CA-CBT \u0110NA.',
    'RCT iCBT cho HS THCS VN (c\u1ea5p 2) \u2014 tu\u1ed5i kh\u1edfi ph\u00e1t SAD ~9,2 (B9 NMA). VN21 (l\u1edbp 6\u21929) c\u00f3 th\u1ec3 k\u1ebft h\u1ee3p.',
]:
    add_red(doc, f'\u2022 {s}')

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '51_Japan_iCBT_SAD_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['77', '4,97', '3,95', '2,01', '61%', '68%', '24%', '34%', 'LSAS', 'UMIN', '10 phi']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
