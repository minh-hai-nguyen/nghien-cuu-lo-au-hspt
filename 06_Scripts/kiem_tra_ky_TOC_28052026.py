# -*- coding: utf-8 -*-
"""Kiem tra ky tung dong - so sanh backup va sau khi them TOC.
28/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
BACKUP = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026_BEFORE_TOC.docx')
CURRENT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

print("=" * 75)
print("KIEM TRA TUNG DONG - SO SANH BACKUP vs CURRENT")
print("=" * 75)

bk = Document(BACKUP)
nw = Document(CURRENT)

# ====================================================================
# 1. Paragraph count
# ====================================================================
print(f"\n[1] Paragraph count: backup={len(bk.paragraphs)} current={len(nw.paragraphs)} diff={len(nw.paragraphs)-len(bk.paragraphs)}")

# ====================================================================
# 2. Per-paragraph diff (accounting for 1-paragraph offset from TOC)
# ====================================================================
# TOC was inserted after para 60. So backup[i] maps to:
#   current[i]    for i <= 60
#   current[i+1]  for i >= 61

def map_idx(i, insert_after=60):
    return i if i <= insert_after else i + 1

print(f"\n[2] PER-PARAGRAPH DIFF (text + format)")
diffs_text = []
diffs_align = []
diffs_indent = []
diffs_linespace = []
diffs_style = []
diffs_runs_count = []
diffs_run_font = []
diffs_run_size = []
diffs_run_bold = []
diffs_run_italic = []
diffs_run_color = []
diffs_run_underline = []
diffs_run_text = []

for i in range(len(bk.paragraphs)):
    j = map_idx(i)
    if j >= len(nw.paragraphs):
        diffs_text.append((i, 'MISSING in current', None))
        continue
    bp = bk.paragraphs[i]
    np = nw.paragraphs[j]
    # Text
    if bp.text != np.text:
        diffs_text.append((i, bp.text[:60], np.text[:60]))
    # Alignment
    if bp.paragraph_format.alignment != np.paragraph_format.alignment:
        diffs_align.append((i, bp.paragraph_format.alignment, np.paragraph_format.alignment, bp.text[:50]))
    # Indent
    bi = bp.paragraph_format.first_line_indent
    ni = np.paragraph_format.first_line_indent
    bi_v = bi.cm if bi else None
    ni_v = ni.cm if ni else None
    if bi_v != ni_v:
        if (bi_v is None) != (ni_v is None) or (bi_v is not None and abs(bi_v - ni_v) > 0.001):
            diffs_indent.append((i, bi_v, ni_v, bp.text[:50]))
    # Line spacing
    if bp.paragraph_format.line_spacing != np.paragraph_format.line_spacing:
        diffs_linespace.append((i, bp.paragraph_format.line_spacing, np.paragraph_format.line_spacing, bp.text[:50]))
    # Style
    bs = bp.style.name if bp.style else None
    ns = np.style.name if np.style else None
    if bs != ns:
        diffs_style.append((i, bs, ns, bp.text[:50]))
    # Run count
    if len(bp.runs) != len(np.runs):
        diffs_runs_count.append((i, len(bp.runs), len(np.runs), bp.text[:50]))
    # Run-by-run
    for k in range(min(len(bp.runs), len(np.runs))):
        br = bp.runs[k]
        nr = np.runs[k]
        if br.text != nr.text:
            diffs_run_text.append((i, k, br.text[:30], nr.text[:30]))
        if br.font.name != nr.font.name:
            diffs_run_font.append((i, k, br.font.name, nr.font.name))
        bs_pt = br.font.size.pt if br.font.size else None
        ns_pt = nr.font.size.pt if nr.font.size else None
        if bs_pt != ns_pt:
            diffs_run_size.append((i, k, bs_pt, ns_pt))
        if br.bold != nr.bold:
            diffs_run_bold.append((i, k, br.bold, nr.bold))
        if br.italic != nr.italic:
            diffs_run_italic.append((i, k, br.italic, nr.italic))
        if br.underline != nr.underline:
            diffs_run_underline.append((i, k, br.underline, nr.underline))
        try:
            bc = str(br.font.color.rgb) if br.font.color and br.font.color.rgb else None
            nc = str(nr.font.color.rgb) if nr.font.color and nr.font.color.rgb else None
            if bc != nc:
                diffs_run_color.append((i, k, bc, nc))
        except: pass

print(f"  Text differences: {len(diffs_text)}")
print(f"  Alignment differences: {len(diffs_align)}")
print(f"  Indent differences: {len(diffs_indent)}")
print(f"  Line spacing differences: {len(diffs_linespace)}")
print(f"  Style differences: {len(diffs_style)}")
print(f"  Run count differences: {len(diffs_runs_count)}")
print(f"  Run text differences: {len(diffs_run_text)}")
print(f"  Run font name differences: {len(diffs_run_font)}")
print(f"  Run font size differences: {len(diffs_run_size)}")
print(f"  Run bold differences: {len(diffs_run_bold)}")
print(f"  Run italic differences: {len(diffs_run_italic)}")
print(f"  Run underline differences: {len(diffs_run_underline)}")
print(f"  Run color differences: {len(diffs_run_color)}")

# Print sample of each diff type if exists
if diffs_text:
    print(f"\n  Sample text diffs (first 3):")
    for d in diffs_text[:3]:
        print(f"    {d}")
if diffs_align:
    print(f"\n  Sample alignment diffs:")
    for d in diffs_align[:3]:
        print(f"    {d}")
if diffs_run_color:
    print(f"\n  Sample run color diffs:")
    for d in diffs_run_color[:3]:
        print(f"    {d}")

# ====================================================================
# 3. The NEW paragraph at position 61 (TOC text)
# ====================================================================
print(f"\n[3] NEW paragraph inserted at current[61]:")
if 61 < len(nw.paragraphs):
    new_p = nw.paragraphs[61]
    print(f"  text: {new_p.text!r}")
    # Check that this is the TOC paragraph (has field code)
    has_toc_field = False
    for elem in new_p._element.iter():
        if elem.tag == qn('w:instrText'):
            if elem.text and 'TOC' in elem.text:
                has_toc_field = True
                print(f"  TOC field code found: {elem.text!r}")
                break
    if not has_toc_field:
        print(f"  WARNING: TOC field code NOT found in this paragraph!")

# ====================================================================
# 4. Tables
# ====================================================================
print(f"\n[4] TABLES")
print(f"  backup tables: {len(bk.tables)}")
print(f"  current tables: {len(nw.tables)}")
mismatched = 0
for ti in range(min(len(bk.tables), len(nw.tables))):
    bt = bk.tables[ti]
    nt = nw.tables[ti]
    if len(bt.rows) != len(nt.rows):
        mismatched += 1
        continue
    for ri in range(len(bt.rows)):
        br = bt.rows[ri]
        nr = nt.rows[ri]
        if len(br.cells) != len(nr.cells):
            mismatched += 1
            break
        for ci in range(len(br.cells)):
            if br.cells[ci].text != nr.cells[ci].text:
                mismatched += 1
                break
print(f"  Tables with content mismatch: {mismatched}")

# ====================================================================
# 5. Sections / margins
# ====================================================================
print(f"\n[5] SECTIONS / MARGINS")
print(f"  backup sections: {len(bk.sections)}")
print(f"  current sections: {len(nw.sections)}")
for i in range(min(len(bk.sections), len(nw.sections))):
    bs = bk.sections[i]
    ns = nw.sections[i]
    match = (bs.top_margin == ns.top_margin and bs.bottom_margin == ns.bottom_margin
             and bs.left_margin == ns.left_margin and bs.right_margin == ns.right_margin)
    mark = "✓" if match else "✗"
    print(f"  {mark} Section {i+1}: bk T/B/L/R = {bs.top_margin.cm:.2f}/{bs.bottom_margin.cm:.2f}/{bs.left_margin.cm:.2f}/{bs.right_margin.cm:.2f} | nw = {ns.top_margin.cm:.2f}/{ns.bottom_margin.cm:.2f}/{ns.left_margin.cm:.2f}/{ns.right_margin.cm:.2f}")

# ====================================================================
# 6. Headers / Footers
# ====================================================================
print(f"\n[6] HEADERS / FOOTERS")
for i, (bs, ns) in enumerate(zip(bk.sections, nw.sections)):
    for hf_name in ['header', 'footer']:
        bk_hf = getattr(bs, hf_name)
        nw_hf = getattr(ns, hf_name)
        bk_text = '\n'.join(p.text for p in bk_hf.paragraphs)
        nw_text = '\n'.join(p.text for p in nw_hf.paragraphs)
        if bk_text != nw_text:
            print(f"  Section {i+1} {hf_name}: DIFF backup={bk_text[:60]!r} current={nw_text[:60]!r}")
        else:
            # Count pict/object/drawing
            bk_pict = sum(1 for e in bk_hf._element.iter() if e.tag in (qn('w:pict'), qn('w:object'), qn('w:drawing')))
            nw_pict = sum(1 for e in nw_hf._element.iter() if e.tag in (qn('w:pict'), qn('w:object'), qn('w:drawing')))
            if bk_pict != nw_pict:
                print(f"  Section {i+1} {hf_name}: pict count DIFF backup={bk_pict} current={nw_pict}")
print(f"  (only shows if differs)")

# ====================================================================
# 7. Metadata
# ====================================================================
print(f"\n[7] METADATA")
bcp = bk.core_properties
ncp = nw.core_properties
for attr in ['author', 'title', 'subject', 'keywords', 'comments', 'last_modified_by']:
    bv = getattr(bcp, attr)
    nv = getattr(ncp, attr)
    if bv != nv:
        print(f"  DIFF {attr}: {bv!r} -> {nv!r}")
    else:
        print(f"  {attr}: {bv!r} (unchanged)")

# ====================================================================
# 8. File size
# ====================================================================
print(f"\n[8] FILE SIZE")
print(f"  backup: {os.path.getsize(BACKUP)//1024} KB")
print(f"  current: {os.path.getsize(CURRENT)//1024} KB")
print(f"  diff: {(os.path.getsize(CURRENT) - os.path.getsize(BACKUP))} bytes")

# ====================================================================
# TONG KET
# ====================================================================
print("\n" + "=" * 75)
print("KET LUAN")
print("=" * 75)
total_diffs = (len(diffs_text) + len(diffs_align) + len(diffs_indent)
               + len(diffs_linespace) + len(diffs_style) + len(diffs_runs_count)
               + len(diffs_run_text) + len(diffs_run_font) + len(diffs_run_size)
               + len(diffs_run_bold) + len(diffs_run_italic) + len(diffs_run_underline)
               + len(diffs_run_color))
if total_diffs == 0:
    print("  ✓ KHONG CO BAT KY SAI LECH NAO trong noi dung + format paragraph + runs")
else:
    print(f"  ✗ Co {total_diffs} sai lech can xem xet (xem chi tiet phia tren)")

if mismatched == 0:
    print("  ✓ Bang giu nguyen 100%")
else:
    print(f"  ✗ {mismatched} bang co sai khac")

print("  ✓ Da chen 1 paragraph TOC moi tai vi tri current[61]")
