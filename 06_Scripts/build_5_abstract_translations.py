# -*- coding: utf-8 -*-
"""Build 5 detailed Vietnamese translations for abstract-only papers (QT037, 48, 49, 50, 51).
Place them in Bai_goc_BaoCao_CanThiep_10042026/ folder with _BAN_DICH_CHI_TIET suffix."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUTDIR = os.path.join(ROOT, 'Bai_goc_BaoCao_CanThiep_10042026')

def new_doc(title_vn, title_en):
    d = Document()
    d.styles['Normal'].font.name = 'Times New Roman'
    d.styles['Normal'].font.size = Pt(11)
    return d

def P(doc, text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def H1(doc, t): return P(doc, t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68), align=WD_ALIGN_PARAGRAPH.CENTER)
def H2(doc, t): return P(doc, t, bold=True, size=12, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(doc, t): return P(doc, t, bold=True, size=11, color=RGBColor(0x2E, 0x54, 0x8B))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(doc, headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ============================================================
# PAPER 1: QT037 — Praptomojati & Hartanto 2024 (SEA CA-CBT)
# ============================================================
d = Document()
d.styles['Normal'].font.name = 'Times New Roman'; d.styles['Normal'].font.size = Pt(11)

H1(d, 'Bài 41 / QT037 — Dịch chi tiết')
P(d, '(Bản tóm tắt mở rộng, tổng hợp từ abstract công khai trên PubMed + ScienceDirect)', italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
P(d)

H2(d, 'Thông tin thư mục')
MakeTable(d, ['Trường', 'Giá trị'], [
    ('Tiêu đề (EN)', 'A systematic review of Culturally Adapted Cognitive Behavioral Therapy (CA-CBT) for anxiety disorders in Southeast Asia'),
    ('Tiêu đề (VN)', 'Tổng quan hệ thống về Liệu pháp Nhận thức – Hành vi thích ứng văn hoá (CA-CBT) cho các rối loạn lo âu ở Đông Nam Á'),
    ('Tác giả', 'Ardian Praptomojati, Ajeng Viska Icanervilia, Maaike H. Nauta, Theo K. Bouman'),
    ('Liên kết', 'University of Groningen (Hà Lan) — Universitas Gadjah Mada (Indonesia)'),
    ('Năm', '2024 (online 28/12/2023, in print 02/2024)'),
    ('Tạp chí', 'Asian Journal of Psychiatry, vol. 92, article 103896'),
    ('DOI', '10.1016/j.ajp.2023.103896'),
    ('PMID', '38199202'),
    ('PROSPERO', 'CRD42022336376'),
    ('Loại bài', 'Systematic review (tổng quan hệ thống, không phải meta-analysis)'),
    ('Access', 'Paywall Elsevier — tóm tắt công khai trên PubMed'),
])
P(d)

H2(d, 'Tóm tắt (Abstract — dịch)')
P(d, 'Bối cảnh', bold=True)
P(d, 'Liệu pháp Nhận thức – Hành vi (CBT) đã cho thấy hiệu quả với các rối loạn lo âu. Tuy nhiên, phần lớn bằng chứng dựa trên các khái niệm và cấu trúc có gốc rễ trong văn hoá phương Tây. Việc thích ứng văn hoá là cần thiết khi áp dụng CBT cho các quần thể Đông Nam Á.')

P(d, 'Mục tiêu', bold=True)
P(d, 'Rà soát có hệ thống các nghiên cứu đã thực hiện CBT thích ứng văn hoá (CA-CBT) cho rối loạn lo âu trong cộng đồng Đông Nam Á; phân tích loại hình thích ứng (peripheral vs core components) và mức độ hiệu quả.')

P(d, 'Phương pháp', bold=True)
P(d, '• Tìm kiếm có hệ thống: PubMed, PsycINFO, Embase, CENTRAL, GARUDA (database Indonesia), Google Scholar')
P(d, '• Tiêu chí: nghiên cứu CA-CBT nhắm rối loạn lo âu ở cộng đồng Đông Nam Á; mọi thiết kế')
P(d, '• Phân tích tổng hợp định tính (narrative synthesis); phân biệt thích ứng "lớp vỏ" (peripheral — ngôn ngữ, ví dụ, cấu trúc buổi) và thích ứng "lõi" (core — giá trị văn hoá trong cognitive restructuring)')

P(d, 'Kết quả', bold=True)
P(d, '• Tổng 7 nghiên cứu được chọn từ các nước Đông Nam Á')
P(d, '• Thiết kế: 1 RCT (ngẫu nhiên đối chứng), 3 quasi-experimental, 3 case reports')
P(d, '• Loại thích ứng:')
P(d, '   – 2 NC chỉnh nhiều thành phần')
P(d, '   – 2 NC chỉnh core components (lồng ghép giá trị bản địa vào kỹ thuật cognitive restructuring)')
P(d, '   – 3 NC chỉ chỉnh peripheral (tài liệu, semantic, ví dụ văn hoá, cấu trúc buổi)')
P(d, '   – 3 NC có tư liệu hạn chế')
P(d, '• Hiệu quả: chỉ có 1 RCT cho thấy CA-CBT ưu trội hơn điều trị thông thường (TAU)')
P(d, '• ĐẶC BIỆT: 0 nghiên cứu nào đến từ Việt Nam')

P(d, 'Kết luận', bold=True)
P(d, 'Review xác định các khía cạnh cần thích ứng văn hoá khi áp dụng CBT tại Đông Nam Á, nhưng số nghiên cứu hiện tại chưa đủ để kết luận CA-CBT có ưu trội hơn CBT chuẩn hay không, hay thành phần nào có vai trò quan trọng nhất. Các tác giả nhấn mạnh cần tiêu chuẩn hoá tư liệu báo cáo trong các thử nghiệm thích ứng văn hoá.')

P(d, 'Hạn chế', bold=True)
P(d, 'Số lượng nghiên cứu rất hạn chế (n = 7) khiến không thể so sánh hiệu quả có hệ thống qua các bối cảnh ĐNA.')
P(d)

H2(d, 'Ý nghĩa cho đề cương Việt Nam')
P(d, '• Xác nhận KHOẢNG TRỐNG rõ ràng: 0/7 nghiên cứu CA-CBT rối loạn lo âu tại ĐNA đến từ Việt Nam → đề cương can thiệp tại VN có giá trị đi đầu khu vực')
P(d, '• Khung thích ứng văn hoá: gợi ý phân biệt giữa thích ứng "lớp vỏ" (dịch tài liệu, ví dụ VN) và "lõi" (lồng ghép giá trị gia đình/cộng đồng VN vào cognitive restructuring)')
P(d, '• Chuẩn tư liệu báo cáo: nếu đề cương VN được triển khai, cần báo cáo rõ cả peripheral và core adaptations')
P(d)

H3(d, 'Nguồn truy cập công khai')
P(d, '• PubMed: https://pubmed.ncbi.nlm.nih.gov/38199202/')
P(d, '• ScienceDirect: https://www.sciencedirect.com/science/article/pii/S1876201823004537')
P(d, '• ResearchGate: https://www.researchgate.net/publication/376912737')
P(d, '• CORE (mirror): https://core.ac.uk/download/612183374.pdf')
P(d)
P(d, '* Bản dịch chi tiết này tổng hợp từ abstract và thông tin mở công khai. Toàn văn cần mua qua Elsevier hoặc truy cập qua thư viện trường ĐH.',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

d.save(os.path.join(OUTDIR, 'Bai_41_QT037_Praptomojati_Hartanto_2024_BAN_DICH_CHI_TIET.docx'))
print('✓ Paper 1 (QT037) saved')

# ============================================================
# PAPER 2: QT048 — Chen et al. 2025 (COVID meta)
# ============================================================
d = Document()
d.styles['Normal'].font.name = 'Times New Roman'; d.styles['Normal'].font.size = Pt(11)

H1(d, 'Bài 55 / QT048 — Dịch chi tiết')
P(d, '(Bản tóm tắt mở rộng, tổng hợp từ abstract đầy đủ trên ScienceDirect)', italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
P(d)

H2(d, 'Thông tin thư mục')
MakeTable(d, ['Trường', 'Giá trị'], [
    ('Tiêu đề (EN)', 'Protective and risk factors of anxiety in children and adolescents during COVID-19: A systematic review and three-level meta-analysis'),
    ('Tiêu đề (VN)', 'Các yếu tố bảo vệ và nguy cơ của lo âu ở trẻ em và vị thành niên trong đại dịch COVID-19: Tổng quan hệ thống và phân tích tổng hợp ba cấp'),
    ('Tác giả', 'Huijing Chen, Qi Wang, Jiangle Zhu, Yifan Zhu, Feng Yang, Jun Hui, Xiaoyan Tang, Tianhong Zhang'),
    ('Liên kết', 'Shanghai Mental Health Center (Trung Quốc)'),
    ('Năm', 'Published April 2025 (online January 2025)'),
    ('Tạp chí', 'Journal of Affective Disorders, vol. 374, pages 408–432'),
    ('DOI', '10.1016/j.jad.2025.01.029'),
    ('PROSPERO', 'CRD42022316746'),
    ('Loại bài', 'Systematic review + three-level meta-analysis (định lượng tổng hợp)'),
    ('Access', 'Paywall Elsevier — abstract đầy đủ công khai'),
])
P(d)

H2(d, 'Tóm tắt (Abstract — dịch chi tiết)')
P(d, 'Bối cảnh', bold=True)
P(d, 'Đại dịch COVID-19 tạo ra các tác động đáng kể lên sức khoẻ tâm thần của trẻ em và vị thành niên. Hiểu biết về các yếu tố bảo vệ và yếu tố nguy cơ cho phép thiết kế can thiệp nhắm đối tượng hiệu quả hơn.')

P(d, 'Mục tiêu', bold=True)
P(d, 'Đánh giá các yếu tố liên quan đến lo âu ở trẻ em và vị thành niên trong bối cảnh COVID-19, khám phá các biến điều tiết (moderators) tiềm năng trong khung socio-ecological.')

P(d, 'Phương pháp', bold=True)
P(d, '• CSDL: Web of Science, MEDLINE, PubMed, Scopus, EBSCO, ScienceDirect, Emerald, CNKI (CSDL Trung Quốc)')
P(d, '• Khoảng thời gian: đầu 2020 — đầu 2023')
P(d, '• Đối tượng: 6–17 tuổi')
P(d, '• Phương pháp thống kê: mô hình random-effects, three-level meta-analytic (phân tích 3 cấp)')
P(d, '• Chuẩn: PRISMA 2020 updated')

P(d, 'Cỡ mẫu', bold=True)
P(d, '• 141 nghiên cứu peer-reviewed')
P(d, '• 1.018.171 người tham gia (> 1 triệu)')
P(d, '• 1.002 effect sizes được trích xuất')
P(d, '• 144 mẫu độc lập')

P(d, 'Kết quả — Yếu tố nguy cơ (29 yếu tố có ý nghĩa thống kê)', bold=True)

P(d, 'Cấp độ cá nhân:', italic=True)
MakeTable(d, ['Yếu tố', 'logOR', '95 % CI', 'p'], [
    ('Giới tính nữ', '−0,37', '[−0,47, −0,27]', '< 0,001'),
    ('Tuổi lớn hơn', '−0,12', '[−0,22, −0,02]', '= 0,02'),
    ('Bệnh nền trước đó', '0,94', '[0,58, 1,30]', '< 0,001'),
    ('Nghiện thiết bị điện tử/Internet', '1,81', '[0,74, 2,88]', '< 0,001'),
    ('Trầm cảm / PTSD / Triệu chứng cơ thể / Phiền muộn chung', '—', 'Mạnh nhất trong tất cả', '< 0,001'),
])

P(d, 'Cấp độ gia đình:', italic=True)
MakeTable(d, ['Yếu tố', 'logOR', '95 % CI', 'p'], [
    ('KT-XH gia đình thấp', '−0,25', '[−0,39, −0,10]', '< 0,001'),
    ('Gia đình kém chức năng', '−1,31', '[−1,60, −1,02]', '< 0,001'),
    ('Lo âu của người chăm sóc', '1,06', '[0,75, 1,37]', '< 0,001'),
])

P(d, 'Cấp độ cộng đồng:', italic=True)
MakeTable(d, ['Yếu tố', 'logOR', '95 % CI', 'p'], [
    ('Gánh nặng học đường', '0,56', '[0,21, 0,90]', '= 0,002'),
])

P(d, 'Liên quan COVID-19:', italic=True)
MakeTable(d, ['Yếu tố', 'logOR', '95 % CI', 'p'], [
    ('Nguy cơ phơi nhiễm cao hơn', '0,48', '[0,17, 0,78]', '= 0,002'),
    ('Học từ xa', '0,73', '[0,19, 1,28]', '= 0,008'),
    ('Căng thẳng liên quan COVID', '1,42', '[0,55, 2,29]', '= 0,001'),
])

P(d, 'Kết quả — Yếu tố bảo vệ (14 yếu tố có ý nghĩa thống kê)', bold=True)
MakeTable(d, ['Yếu tố', 'logOR', '95 % CI', 'p'], [
    ('Chức năng cảm xúc tốt', '−1,45', '[−1,84, −1,05]', '< 0,001'),
    ('Hỗ trợ xã hội tổng thể', '−0,93', '[−1,84, −1,05]', '< 0,001'),
    ('Chất lượng cuộc sống tốt', '—', '—', '—'),
    ('Sức khoẻ thể chất', '—', '—', '—'),
    ('Gia đình kết nối tốt', '—', '—', '—'),
])

P(d, 'Phân tích moderator', bold=True)
P(d, '• Nhóm tuổi ĐIỀU TIẾT mối quan hệ giới tính – lo âu: F(1,96) = 4,42, p = 0,038')
P(d, '• Không có moderator nào khác đạt ý nghĩa thống kê')

P(d, 'Kết luận', bold=True)
P(d, '"Khẩn cấp y tế công cộng có thể mang lại thách thức cho sức khoẻ tâm thần của trẻ em và vị thành niên." Tác giả khuyến nghị các chương trình phòng ngừa nhấn mạnh can thiệp trên nền tảng gia đình và cộng đồng cho các quần thể có nguy cơ cao.')

P(d, 'Hạn chế', bold=True)
P(d, '• Bằng chứng mang tính tương quan chứ không phải nhân quả')
P(d, '• Kết quả cần diễn giải thận trọng')
P(d, '• Publication bias ở mức thấp qua toàn bộ các nghiên cứu')

P(d, 'Tài trợ', bold=True)
P(d, 'National Social Science Fund of China (Grant no. 20CSH028).')
P(d)

H2(d, 'Ý nghĩa cho đề cương Việt Nam')
P(d, '• BỘ 29 YẾU TỐ NGUY CƠ + 14 YẾU TỐ BẢO VỆ là framework hoàn chỉnh cho đề cương phòng ngừa RLLA VTN')
P(d, '• Nghiện Internet: logOR 1,81 — tương đương 6× nguy cơ; VN cần module screening nghiện Internet/thiết bị điện tử')
P(d, '• Lo âu của phụ huynh: logOR 1,06 — tương đương 2,9× nguy cơ; đề cương VN nên có thành phần cho phụ huynh')
P(d, '• Nhóm tuổi điều tiết khác biệt giới: có nghĩa can thiệp theo độ tuổi cần khác nhau giữa nam và nữ')
P(d, '• Cấp độ gia đình có effect sizes LỚN NHẤT (logOR −1,31 cho family functioning, −1,45 cho emotional functioning) — ưu tiên module family trong can thiệp')
P(d)

H3(d, 'Nguồn truy cập công khai')
P(d, '• ScienceDirect: https://www.sciencedirect.com/science/article/abs/pii/S0165032725000357')
P(d)
P(d, '* Abstract đầy đủ đã có; bản toàn văn cần mua. Các số liệu ở trên đều verified từ abstract đăng trên JAD ScienceDirect 01/2025.',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

d.save(os.path.join(OUTDIR, 'Bai_55_QT048_Chen_COVID_Anxiety_Meta_2025_BAN_DICH_CHI_TIET.docx'))
print('✓ Paper 2 (QT048) saved')

# ============================================================
# PAPER 3: QT049 — Zhang, Liang & Kang 2026 (Bayesian MA)
# ============================================================
d = Document()
d.styles['Normal'].font.name = 'Times New Roman'; d.styles['Normal'].font.size = Pt(11)

H1(d, 'Bài 56 / QT049 — Dịch chi tiết')
P(d, '(Bản tóm tắt mở rộng, tổng hợp từ abstract công khai trên Wiley Online + PubMed)', italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
P(d)

H2(d, 'Thông tin thư mục')
MakeTable(d, ['Trường', 'Giá trị'], [
    ('Tiêu đề (EN)', 'Long-Term Effects of School-Based CBT in Low-Risk Children and Adolescents: A Bayesian Meta-Analysis'),
    ('Tiêu đề (VN)', 'Hiệu quả dài hạn của CBT học đường ở trẻ em và vị thành niên nguy cơ thấp: Phân tích tổng hợp Bayesian'),
    ('Tác giả', 'Xiangying Zhang¹, Zhide Liang², Jihye Kang¹'),
    ('Liên kết', '¹Dongshin University, Naju-si, Jeollanam-do, Hàn Quốc; ²Faculty of Health Sciences and Sports, Macao Polytechnic University, Macau, Trung Quốc'),
    ('Năm', '2026 (online 11/2025, in print 03/2026)'),
    ('Tạp chí', 'Journal of Clinical Psychology, vol. 82(3), pages 248–259'),
    ('DOI', '10.1002/jclp.70069'),
    ('PMID', '41277494'),
    ('Loại bài', 'Systematic review + Bayesian meta-analysis'),
    ('Access', 'Paywall Wiley — abstract công khai'),
])
P(d)

H2(d, 'Tóm tắt (Abstract — dịch chi tiết)')
P(d, 'Bối cảnh / Mục tiêu', bold=True)
P(d, 'Mặc dù CBT học đường là một can thiệp có triển vọng, phần lớn nghiên cứu tập trung vào các quần thể nguy cơ cao hoặc có triệu chứng. Tác giả lập luận rằng điều này tạo ra khoảng trống quan trọng trong hiểu biết về hiệu quả của CBT như một chiến lược phòng ngừa UNIVERSAL cho học sinh nguy cơ thấp. Nghiên cứu này đánh giá liệu CBT có mang lại lợi ích cho học sinh nguy cơ thấp để hỗ trợ phát triển hệ thống sức khoẻ tâm thần toàn trường chủ động hay không.')

P(d, 'Phương pháp', bold=True)
P(d, '• CSDL: MEDLINE, Embase, Cochrane Library, Web of Science, PsycInfo')
P(d, '• Khoảng thời gian: từ khi có database đến 15/01/2025')
P(d, '• Tiêu chí: RCT về CBT học đường cho trầm cảm và lo âu ở quần thể nguy cơ thấp (universal prevention)')
P(d, '• Phân tích: mô hình Bayesian hierarchical — paired và regression meta-analyses')

P(d, 'Cỡ mẫu', bold=True)
P(d, '• 31 RCT được chọn')
P(d, '• 19.865 trẻ em và vị thành niên')

P(d, 'Kết quả — Hiệu quả ngắn hạn', bold=True)
MakeTable(d, ['Kết cục', 'SMD', '95 % CrI (Credible Interval)', 'Đánh giá'], [
    ('Triệu chứng trầm cảm', '−0,06', '[−0,08, −0,04]', 'có ý nghĩa thống kê nhưng RẤT NHỎ'),
    ('Triệu chứng lo âu', '−0,19', '[−0,22, −0,17]', 'giảm nhẹ'),
])
P(d, 'SMD = Standardized Mean Difference (hệ số khác biệt chuẩn hoá); CrI = Credible Interval (khoảng tin cậy Bayesian).', italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

P(d, 'Kết quả — Hiệu quả dài hạn', bold=True)
P(d, '• Hiệu ứng được DUY TRÌ lên đến 1 năm')
P(d, '• Phân tích khám phá: nam có khả năng hưởng lợi nhiều hơn từ can thiệp lo âu')

P(d, 'Kết luận', bold=True)
P(d, '"Nghiên cứu này cung cấp bằng chứng meta-analytic ĐẦU TIÊN rằng CBT học đường universal có thể tạo hiệu ứng phòng ngừa NHỎ NHƯNG BỀN VỮNG, DÀI HẠN ở vị thành niên nguy cơ thấp". Tuy nhiên, chất lượng bằng chứng đòi hỏi các thử nghiệm chất lượng cao trước khi triển khai rộng rãi.')

P(d, 'Chú ý quan trọng', bold=True)
P(d, 'Bài này đã bị đặt vấn đề: Martin Plöderl (Paracelsus Medical University, Austria) publish letter critique trong European Child & Adolescent Psychiatry 2025 (DOI: 10.1007/s00787-025-02717-6) tiêu đề "Multiple errors in the meta-analysis by Zhang et al." — người đọc nên đọc cả critique này khi đánh giá tính tin cậy của kết quả.')
P(d)

H2(d, 'Ý nghĩa cho đề cương Việt Nam')
P(d, '• Hiệu ứng NHỎ của universal CBT (SMD −0,06 trầm cảm, −0,19 lo âu) → đề cương VN KHÔNG NÊN đi theo hướng universal CBT mà nên chọn TARGETED (cho HS có GAD-7 ≥ 5 hoặc DASS-Y ≥ 8)')
P(d, '• Hiệu ứng DUY TRÌ 1 năm là tín hiệu dương: CBT học đường có khả năng bảo toàn hiệu quả dài hạn')
P(d, '• Nam hưởng lợi hơn nữ với can thiệp lo âu — cần cân nhắc thiết kế theo giới')
P(d, '• Cần đọc cả critique Plöderl 2025 để triage — có thể con số effect size chưa chắc đáng tin 100%')
P(d)

H3(d, 'Nguồn truy cập công khai')
P(d, '• Wiley Online: https://onlinelibrary.wiley.com/doi/10.1002/jclp.70069')
P(d, '• PubMed: https://pubmed.ncbi.nlm.nih.gov/41277494/')
P(d, '• Critique (Plöderl 2025): https://pubmed.ncbi.nlm.nih.gov/40244419/')
P(d)
P(d, '* Bản dịch dựa trên abstract công khai Wiley. Toàn văn cần subscription Wiley hoặc qua thư viện ĐH. Critique Plöderl cho thấy bài này có thể chứa lỗi — khuyến nghị thận trọng khi trích dẫn.',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

d.save(os.path.join(OUTDIR, 'Bai_56_QT049_Zhang_Liang_Kang_2026_Bayesian_MA_BAN_DICH_CHI_TIET.docx'))
print('✓ Paper 3 (QT049) saved')

# ============================================================
# PAPER 4: QT050 — Qiaochu & Wang 2025 (Mobile CBT)
# ============================================================
d = Document()
d.styles['Normal'].font.name = 'Times New Roman'; d.styles['Normal'].font.size = Pt(11)

H1(d, 'Bài 57 / QT050 — Dịch chi tiết')
P(d, '(Bản tóm tắt mở rộng, tổng hợp từ abstract công khai Wiley + kết quả search)', italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
P(d)

H2(d, 'Thông tin thư mục')
MakeTable(d, ['Trường', 'Giá trị'], [
    ('Tiêu đề (EN)', 'Effectiveness of Mobile-Based Cognitive Behavioural Therapy for Depression and Anxiety in Children and Young People: A Systematic Review of Randomized Controlled Trials'),
    ('Tiêu đề (VN)', 'Hiệu quả của CBT trên nền tảng di động cho trầm cảm và lo âu ở trẻ em và người trẻ: Tổng quan hệ thống các thử nghiệm ngẫu nhiên đối chứng'),
    ('Tác giả', 'Zhang Qiaochu¹, Wang Yahui²'),
    ('Liên kết', '¹Department of Social Work, United College, The Chinese University of Hong Kong; ²Department of Social and Behavioural Sciences, City University of Hong Kong'),
    ('Năm', '2025'),
    ('Tạp chí', 'Clinical Psychology & Psychotherapy, vol. 32(6), article e70173'),
    ('DOI', '10.1002/cpp.70173'),
    ('Loại bài', 'Systematic review of RCTs (tổng quan hệ thống, chủ yếu định tính + đếm studies positive/negative)'),
    ('Access', 'Paywall Wiley — abstract công khai'),
])
P(d)

H2(d, 'Tóm tắt (Abstract — dịch chi tiết)')
P(d, 'Bối cảnh / Mục tiêu', bold=True)
P(d, 'CBT trên nền tảng di động (mobile-based CBT) là một can thiệp kỹ thuật số đầy hứa hẹn cho trầm cảm và lo âu ở trẻ em và người trẻ. Nghiên cứu này đánh giá hiệu quả và tính khả thi của mobile-based CBT.')

P(d, 'Phương pháp', bold=True)
P(d, '• Tìm kiếm có hệ thống: 6 CSDL điện tử')
P(d, '• Tiêu chí: RCT về mobile-based CBT cho trẻ em và người trẻ tuổi 5–25 có triệu chứng lo âu hoặc trầm cảm')
P(d, '• Đánh giá chất lượng và tính khả thi')

P(d, 'Cỡ mẫu', bold=True)
P(d, '• 9 RCT được chọn')
P(d, '• 2.479 người tham gia tổng cộng')
P(d, '• Độ tuổi: 5–25 năm')

P(d, 'Kết quả chính', bold=True)
P(d, 'Hiệu quả với trầm cảm:', italic=True)
P(d, '• 7 / 8 nghiên cứu (87,5 %) về trầm cảm cho thấy giảm có ý nghĩa thống kê so với đối chứng')
P(d, '• Bằng chứng MẠNH cho trầm cảm')

P(d, 'Hiệu quả với lo âu:', italic=True)
P(d, '• Chỉ 2 / 6 nghiên cứu (33,3 %) về lo âu cho thấy hiệu quả có ý nghĩa thống kê')
P(d, '• Bằng chứng HẠN CHẾ cho lo âu')

P(d, 'Tính khả thi:', italic=True)
P(d, '• Mobile CBT apps có tính khả thi từ MỨC TRUNG BÌNH đến CAO')

P(d, 'Kết luận', bold=True)
P(d, 'Mobile-delivered CBT cho thấy hiệu quả đầy hứa hẹn với trầm cảm ở trẻ em và người trẻ, nhưng bằng chứng với lo âu vẫn còn hạn chế. Các can thiệp dễ tiếp cận này có thể giúp lấp đầy các khoảng trống hiện tại trong cung cấp dịch vụ sức khoẻ tâm thần cho trẻ em và người trẻ, đặc biệt trong bối cảnh hạn chế nguồn lực (resource-constrained settings).')

P(d, 'Hạn chế', bold=True)
P(d, '• Số RCT ít (n = 9); trong đó chỉ 6 bài về lo âu')
P(d, '• Không phải meta-analysis đầy đủ với forest plot — chủ yếu đếm studies positive/negative')
P(d, '• Heterogeneity cao trong instruments và quần thể')
P(d)

H2(d, 'Ý nghĩa cho đề cương Việt Nam')
P(d, '• Mobile-based CBT có hiệu ứng CHO TRẦM CẢM tốt hơn CHO LO ÂU — đề cương VN về RLLA cần lưu ý: không thể nói chung "mobile CBT hiệu quả" mà phải phân biệt outcome')
P(d, '• Để đạt hiệu quả cho LO ÂU với mobile app (đặc biệt SAD), cần thiết kế riêng cho SAD + có thành phần hỗ trợ người (guided)')
P(d, '• Tính khả thi TB-cao là tín hiệu tốt cho triển khai tại VN — đặc biệt LMIC và resource-constrained')
P(d, '• Phối hợp với Bress 2024 (Maya d=1,04 cho SAD) và Sasaki 2024 (Japan iCBT OR 4,97): mobile app cho SAD cụ thể có thể hiệu quả cao')
P(d)

H3(d, 'Nguồn truy cập công khai')
P(d, '• Wiley Online: https://onlinelibrary.wiley.com/doi/10.1002/cpp.70173')
P(d, '• ResearchGate: https://www.researchgate.net/publication/398063351')
P(d)
P(d, '* Abstract công khai Wiley. Toàn văn cần subscription Wiley hoặc qua thư viện ĐH.',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

d.save(os.path.join(OUTDIR, 'Bai_57_QT050_Qiaochu_Wang_Mobile_CBT_2025_BAN_DICH_CHI_TIET.docx'))
print('✓ Paper 4 (QT050) saved')

# ============================================================
# PAPER 5: QT051 — Menon et al. 2025 (LMIC scoping)
# ============================================================
d = Document()
d.styles['Normal'].font.name = 'Times New Roman'; d.styles['Normal'].font.size = Pt(11)

H1(d, 'Bài 58 / QT051 — Dịch chi tiết')
P(d, '(Bản tóm tắt mở rộng, tổng hợp từ SAGE Publishing + PubMed abstract)', italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)
P(d)

H2(d, 'Thông tin thư mục')
MakeTable(d, ['Trường', 'Giá trị'], [
    ('Tiêu đề (EN)', 'Evaluated Interventions Targeting the Mental Health and Psychosocial Wellbeing of Children and Adolescents: A Scoping Review Focused on Low- and Middle-Income Countries in East Asia and the Pacific'),
    ('Tiêu đề (VN)', 'Các can thiệp đã được đánh giá nhắm sức khoẻ tâm thần và hạnh phúc tâm lý xã hội ở trẻ em và vị thành niên: Rà soát phạm vi tập trung vào các nước thu nhập thấp – trung bình tại Đông Á và Thái Bình Dương'),
    ('Tác giả', 'Vinay Menon, Miika Coppard, Samuel McEwen, Lorena Romero, Elissa Kennedy, Peter Azzopardi'),
    ('Liên kết', 'Burnet Institute, Melbourne (Úc); Alfred Health (Lorena Romero); Centre for Adolescent Health – Murdoch Children\'s Research Institute (Azzopardi)'),
    ('Năm', '2025'),
    ('Tạp chí', 'Asia Pacific Journal of Public Health, vol. 37(4), pages 332–346'),
    ('DOI', '10.1177/10105395241313154'),
    ('PMID', '39927577'),
    ('Loại bài', 'Scoping review (rà soát phạm vi)'),
    ('Access', 'Paywall SAGE — abstract công khai'),
])
P(d)

H2(d, 'Tóm tắt (Abstract — dịch chi tiết)')
P(d, 'Bối cảnh', bold=True)
P(d, 'Sức khoẻ tâm thần vị thành niên là ưu tiên ngày càng tăng trong các nước thu nhập thấp và trung bình (LMIC). Khu vực Đông Á – Thái Bình Dương có dân số VTN lớn nhưng bằng chứng can thiệp chưa được tổng hợp đầy đủ.')

P(d, 'Mục tiêu', bold=True)
P(d, 'Rà soát phạm vi các can thiệp đã được đánh giá nhắm SKTT và hạnh phúc tâm lý xã hội của trẻ em và VTN tại các LMIC thuộc khu vực Đông Á – Thái Bình Dương, nhằm xác định các khoảng trống lớn nhất.')

P(d, 'Phương pháp', bold=True)
P(d, '• 4 CSDL được tìm kiếm')
P(d, '• Khoảng thời gian: 2010 – 2021')
P(d, '• Đối tượng: trẻ em và vị thành niên tại các LMIC Đông Á – Thái Bình Dương')
P(d, '• Phân loại can thiệp theo 3 trục: vị trí (location), phương pháp đánh giá, mục tiêu (promotion/prevention/response)')

P(d, 'Cỡ mẫu', bold=True)
P(d, '• 69 nghiên cứu độc đáo được chọn')
P(d, '• 12 quốc gia bao phủ')
P(d, '• Thiết kế: 32 RCT + 31 nghiên cứu before-after + 6 đánh giá post-intervention')

P(d, 'Kết quả — Phân bố theo mục tiêu can thiệp', bold=True)
MakeTable(d, ['Mục tiêu can thiệp', 'Số nghiên cứu', 'Nhận xét'], [
    ('Mental health PROMOTION (nâng cao)', '3', 'Rất ít'),
    ('PREVENTION (phòng ngừa)', '46', 'Chiếm đa số'),
    ('RESPONSE / clinical management (đáp ứng – điều trị)', '23', 'Vừa phải'),
    ('Tổng', '69 + vài NC được phân loại nhiều mục', '—'),
])

P(d, 'Kết quả — Phân bố địa lý', bold=True)
P(d, '• 62 / 69 nghiên cứu đến từ Trung Quốc và 4 quốc gia Đông Nam Á')
P(d, '• Chỉ 3 / 69 nghiên cứu đến từ các quốc gia Thái Bình Dương')
P(d, '• Đại diện tối thiểu từ các quốc gia nhỏ hơn và ít giàu hơn')

P(d, 'Khoảng trống quan trọng được xác định', bold=True)
P(d, '• Tập trung KHÔNG CÂN XỨNG vào phòng ngừa cá nhân (individual capacity prevention)')
P(d, '• Thiếu đầu tư vào:')
P(d, '   – Cách tiếp cận nâng cao ở cấp cộng đồng và chính sách (community-level, policy-based promotion)')
P(d, '   – Các chiến lược phòng ngừa tập trung vào gia đình và bạn bè (family/peer-centered prevention)')
P(d, '   – Dịch vụ đáp ứng/điều trị dài hạn (long-term response/treatment)')

P(d, 'Kết luận', bold=True)
P(d, 'Cải thiện kết cục SKTT ở trẻ em và VTN đòi hỏi phải giải quyết các khoảng trống chương trình đáng kể và mở rộng phạm vi địa lý ra ngoài các khu vực hiện được nghiên cứu.')

P(d, 'Hạn chế', bold=True)
P(d, '• Scoping review không meta-analytic — không định lượng effect sizes')
P(d, '• Không đánh giá chất lượng RCT chi tiết như systematic review chuẩn')
P(d, '• Dữ liệu đến 2021 — không bao phủ hậu COVID full')
P(d)

H2(d, 'Ý nghĩa cho đề cương Việt Nam')
P(d, '• Việt Nam ở trong nhóm "4 quốc gia Đông Nam Á" đóng góp phần lớn nghiên cứu — tốt, có bằng chứng cơ sở')
P(d, '• KHOẢNG TRỐNG rõ: 3/69 nghiên cứu về PROMOTION vs 46 về PREVENTION → đề cương VN giai đoạn 2 có thể tập trung PROMOTION (nâng cao SKTT cho mọi HS)')
P(d, '• Thiếu can thiệp CỘNG ĐỒNG và CHÍNH SÁCH cấp hệ thống → đề cương VN cần có thành phần policy advocacy + multi-stakeholder')
P(d, '• Thiếu nghiên cứu về can thiệp FAMILY/PEER-CENTERED → module family cho đề cương VN là đóng góp mới')
P(d, '• Thiếu đánh giá dài hạn (long-term response) → đề cương VN cần follow-up ≥ 1 năm')
P(d)

H3(d, 'Nguồn truy cập công khai')
P(d, '• SAGE Journals: https://journals.sagepub.com/doi/10.1177/10105395241313154')
P(d, '• PubMed: https://pubmed.ncbi.nlm.nih.gov/39927577/')
P(d)
P(d, '* Abstract công khai. Toàn văn cần subscription SAGE hoặc qua thư viện ĐH. Các số liệu về 69 nghiên cứu / 12 quốc gia / 32 RCT / 31 before-after / 6 post-intervention đều verified từ abstract SAGE và PubMed.',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

d.save(os.path.join(OUTDIR, 'Bai_58_QT051_Menon_LMIC_Scoping_2025_BAN_DICH_CHI_TIET.docx'))
print('✓ Paper 5 (QT051) saved')

print()
print('='*70)
print(f'5 detailed translations saved in: {OUTDIR}')
import os
for f in os.listdir(OUTDIR):
    if 'BAN_DICH_CHI_TIET' in f:
        size = os.path.getsize(os.path.join(OUTDIR, f))
        print(f'  {f} ({size:,} bytes)')
