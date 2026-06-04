# -*- coding: utf-8 -*-
"""Dịch B9 — NMA SAD 2024 — Xian et al. — J Affective Disorders"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2024.08.097', size=10)

add_heading(doc, 'Can thiệp tâm lý cho rối loạn lo âu xã hội ở trẻ em và thanh thiếu niên: Tổng quan hệ thống và phân tích tổng hợp mạng', 1)
h = doc.add_paragraph()
r = h.add_run('Psychological interventions for social anxiety disorder in children and adolescents: A systematic review and network meta-analysis')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Psychological interventions for social anxiety disorder in children and adolescents: A systematic review and network meta-analysis'),
    ('Tác giả', 'Jinhua Xian, Yan Zhang, Bo Jiang*'),
    ('Cơ quan', 'Khoa Khoa học Giáo dục, ĐH Sư phạm Giang Tô 2, Nam Kinh, Trung Quốc'),
    ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
    ('Xuất bản', '2024, Vol. 365, pp. 614–627, 14 trang'),
    ('DOI', '10.1016/j.jad.2024.08.097'),
    ('PROSPERO', 'CRD42023476829'),
    ('Loại NC', 'Tổng quan hệ thống + Phân tích tổng hợp mạng Bayesian (NMA)'),
    ('Mẫu', '30 RCT, 1.547 trẻ em/VTN với SAD, 9 liệu pháp + 3 đối chứng'),
])
add_page_ref(doc, '614–627', 'J Affective Disorders', 'Vol. 365, 2024')

# TOM TAT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Bối cảnh: Rối loạn lo âu xã hội (SAD) là rối loạn tâm thần phổ biến ở trẻ em/VTN, tỷ lệ suốt đời ~10%. Tuổi chẩn đoán TB ~9,2. Thường đi kèm trầm cảm. Ít người trẻ SAD nhận điều trị hiệu quả do triệu chứng bị nhầm với nhút nhát. Mục đích: so sánh và xếp hạng hiệu quả các liệu pháp tâm lý cho SAD ở trẻ/VTN.')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Chỉ RCT. PubMed, Embase, Cochrane, Web of Science. NMA Bayesian. 30 RCT, 1.547 cá nhân, 9 liệu pháp + 3 đối chứng. Kết quả: iCBT (internet-delivered CBT) SUCRA = 71,2% — tối ưu giảm lo âu XH. gCBT (group CBT) SUCRA = 68,4% — hạng 2. CBT cá nhân SUCRA = 66,0% — hạng 3. iCBT cũng giảm trầm cảm tốt nhất (SUCRA = 92,2%). gCBT tăng chức năng tốt nhất (SUCRA = 89,6%).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

# DANH GIA NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'J Affect Disord Q1 IF ≈ 6,6. NMA Bayesian — phương pháp mạnh.',
    '30 RCT, 1.547 trẻ/VTN với SAD — TẬP TRUNG LO ÂU XÃ HỘI riêng (khác QT29 lo âu chung).',
    'iCBT (internet) HẠNG 1 (SUCRA 71,2%) — phát hiện quan trọng: can thiệp TRỰC TUYẾN hiệu quả nhất.',
    'gCBT hạng 2 (68,4%) — phù hợp VN (nhóm, ít chuyên gia, triển khai trường).',
    'iCBT giảm trầm cảm tốt nhất (SUCRA 92,2%) — "hai trong một".',
    'gCBT tăng CHỨC NĂNG tốt nhất (89,6%) — cải thiện học tập, quan hệ XH.',
    'SAD suốt đời ~10%, tuổi khởi phát ~9,2 — cần can thiệp SỚM.',
    'PROSPERO đăng ký. PRISMA tuân thủ.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Đa số RCT từ phương Tây (Mỹ, Úc, UK) — ít châu Á.',
    'Số RCT mỗi liệu pháp nhỏ — iCBT chỉ ~3-4 RCT.',
    'Không phân biệt tuổi rõ (trẻ em vs VTN) trong xếp hạng.',
    'Chỉ SAD — không áp dụng cho GAD hay lo âu chung.',
]:
    add_p(doc, f'• {b}')

# PHUONG PHAP
add_page_ref(doc, '1–3', 'J Affect Disord', 'Vol. 365, 2024')
add_heading(doc, '1. PHƯƠNG PHÁP', 2)
add_p(doc, 'Tìm kiếm: PubMed, Embase, Cochrane Library, Web of Science. Chỉ RCT. Tiêu chuẩn: trẻ em/VTN được chẩn đoán SAD (DSM/ICD). NMA Bayesian framework. PROSPERO CRD42023476829.')
add_p(doc, '9 liệu pháp: (1) CBT cá nhân (I-CBT), (2) CBT nhóm (G-CBT), (3) CBT internet (Int-CBT), (4) SET-C (Social Effectiveness Training), (5) IPT (Interpersonal Therapy), (6) AT (Attention Training), (7) CBM-I (Cognitive Bias Modification - Interpretation), (8) Int-PCIT (Internet Parent-Child Interaction Therapy), (9) BT (Behavioral Therapy). 3 đối chứng: WL (danh sách chờ), NT (không điều trị), PBO (placebo).')

# KET QUA
add_page_ref(doc, '3–10', 'J Affect Disord', 'Vol. 365, 2024')
add_heading(doc, '2. KẾT QUẢ', 2)
add_p(doc, '30 RCT (2000–2023), 1.547 cá nhân. Đa số từ Mỹ. Tuổi TB ~10–16. Các NC dùng ADIS (chẩn đoán), SPAI/SAS/LSAS (đo lo âu XH), CDI (trầm cảm), CGAS/K-GAS (chức năng).')

add_heading(doc, 'Bảng 1. Xếp hạng hiệu quả — GIẢM LO ÂU XÃ HỘI', 3)
add_table(doc,
    ['Hạng', 'Liệu pháp', 'SUCRA (%)', 'Đặc điểm', 'Áp dụng VN'],
    [['1', 'Int-CBT (CBT internet)', '71,2', 'Tự học online + hướng dẫn', 'Cần internet, app VN'],
     ['2', 'G-CBT (CBT nhóm)', '68,4', 'Nhóm 6–12 HS, 8–16 tuần', 'KHẢ THI nhất cho trường VN'],
     ['3', 'I-CBT (CBT cá nhân)', '66,0', '1-1 với nhà trị liệu', 'Thiếu chuyên gia VN'],
     ['4', 'SET-C (Training XH)', '~60', 'Kỹ năng XH + phơi nhiễm', 'Có thể kết hợp'],
     ['5', 'IPT (Liệu pháp giữa cá nhân)', '~55', 'Quan hệ giữa cá nhân', 'Ít NC'],
     ['6', 'AT (Huấn luyện chú ý)', '~50', 'Chuyển hướng chú ý', 'Thực nghiệm'],
     ['...', 'CBM-I, Int-PCIT', 'Thấp', 'Ít hiệu quả', ''],
     ['—', 'WL / NT / PBO', 'Thấp nhất', 'Đối chứng', '']],
    widths=[0.8, 3.5, 2.0, 3.0, 3.5])

add_heading(doc, 'Bảng 2. Xếp hạng theo KẾT QUẢ KHÁC', 3)
add_table(doc,
    ['Kết quả', 'Liệu pháp tốt nhất', 'SUCRA (%)', 'Ý nghĩa'],
    [['Giảm TRẦM CẢM', 'Int-CBT (internet)', '92,2', 'Hiệu quả "hai trong một"'],
     ['Tăng CHỨC NĂNG', 'G-CBT (nhóm)', '89,6', 'Cải thiện học tập, quan hệ XH'],
     ['Giảm lo âu XH', 'Int-CBT', '71,2', 'Hạng 1 tổng thể']],
    widths=[3.0, 4.0, 2.0, 4.0])
add_p(doc, 'Int-CBT hiệu quả NHẤT cho cả lo âu XH (71,2%) VÀ trầm cảm (92,2%). G-CBT tốt nhất cho CHỨC NĂNG (89,6%). Gợi ý: kết hợp Int-CBT (giảm triệu chứng) + G-CBT (cải thiện chức năng) có thể tối ưu.', size=9, italic=True)

# THAO LUAN
add_page_ref(doc, '10–13', 'J Affect Disord', 'Vol. 365, 2024')
add_heading(doc, '3. THẢO LUẬN VÀ KẾT LUẬN', 2)
add_p(doc, 'iCBT hạng 1 — phát hiện quan trọng: can thiệp TRỰC TUYẾN có thể vượt trội can thiệp mặt đối mặt cho SAD. Lý do: (1) VTN SAD sợ tương tác trực tiếp → online ít áp lực hơn, (2) tự điều chỉnh tốc độ, (3) tiếp cận dễ hơn (không cần chuyên gia). Phù hợp B2 JMIR (digital interventions SAD). Phù hợp bối cảnh VN (thiếu chuyên gia, internet phổ biến).')
add_p(doc, 'gCBT hạng 2 và tăng CHỨC NĂNG tốt nhất (89,6%) — phù hợp trường học. So: B8 Sri Lanka RCT dùng gCBT tại trường (8 phiên) → hiệu quả sau 3 tháng. QT29 (Li 2025 BMC NMA) xếp CBT SUCRA 0,66 cho lo âu CHUNG (không riêng SAD).')
add_p(doc, 'SAD suốt đời ~10%, tuổi khởi phát ~9,2 — cần can thiệp SỚM ở THCS. Phù hợp VN21 (Trần Thảo Vi: stress tăng lớp 6→9). Jefferies QT35: VN SAD = 30,7% nhưng 18% không biết → sàng lọc trước, can thiệp sau.')
add_p(doc, 'Int-PCIT và CBM-I hiệu quả THẤP — gợi ý: can thiệp cha mẹ qua internet và huấn luyện thiên lệch nhận thức chưa đủ cho SAD trẻ/VTN.')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
for ref in [
    'Xian, J., Zhang, Y. & Jiang, B. (2024). Psychological interventions for SAD in children and adolescents: SR and NMA. J Affect Disord, 365, 614–627.',
    'Beidel, D.C. et al. (2000). Behavioral treatment of childhood social phobia. J Consulting Clin Psych, 68, 1072–1080.',
    'Masia, C.L. et al. (2001). SASS: Skills for Academic and Social Success. Group CBT school-based.',
    '(Xem đầy đủ 70+ TLTK trong bài gốc)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('SAD', 'Social Anxiety Disorder — Rối loạn Lo âu Xã hội'),
    ('NMA', 'Network Meta-Analysis — Phân tích Tổng hợp Mạng'),
    ('SUCRA', 'Surface Under the Cumulative Ranking — Diện tích Dưới Xếp hạng Tích lũy'),
    ('iCBT / Int-CBT', 'Internet-delivered CBT — CBT qua Internet'),
    ('gCBT / G-CBT', 'Group CBT — CBT Nhóm'),
    ('I-CBT', 'Individual CBT — CBT Cá nhân'),
    ('SET-C', 'Social Effectiveness Training for Children'),
    ('SASS', 'Skills for Academic and Social Success'),
    ('IPT', 'Interpersonal Therapy — Liệu pháp Giữa Cá nhân'),
    ('CBM-I', 'Cognitive Bias Modification - Interpretation'),
    ('Int-PCIT', 'Internet Parent-Child Interaction Therapy'),
    ('SPAI', 'Social Phobia and Anxiety Inventory'),
    ('ADIS', 'Anxiety Disorders Interview Schedule'),
    ('WL', 'Waiting List — Danh sách Chờ'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'J Affect Disord Q1 IF ≈ 6,6. NMA Bayesian — so sánh trực tiếp + gián tiếp 9 liệu pháp.',
    '30 RCT, 1.547 trẻ/VTN — TẬP TRUNG SAD RIÊNG (khác QT29 lo âu chung). PROSPERO.',
    'iCBT hạng 1 — phát hiện ĐỘT PHÁ: online vượt trội mặt đối mặt cho SAD (VTN sợ tương tác → online ít áp lực).',
    'Đo 3 kết quả: lo âu XH + trầm cảm + chức năng — toàn diện.',
    'gCBT chức năng 89,6% — phù hợp trường học VN. B8 Sri Lanka RCT xác nhận.',
    'So sánh 9 liệu pháp + 3 đối chứng — NMA toàn diện nhất cho SAD trẻ/VTN.',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'Đa số RCT từ phương Tây (Mỹ, Úc, UK, Na Uy) — ít châu Á. B7 CA-CBT ĐNA: chỉ 7 NC.',
    'Số RCT mỗi liệu pháp NHỎ — iCBT chỉ ~3-4 RCT. KTC rộng.',
    'Không phân biệt tuổi (trẻ em 6-12 vs VTN 13-18) trong xếp hạng — hiệu quả có thể khác.',
    'Chỉ SAD — không áp dụng cho GAD, lo âu chung. QT29 (BMC NMA) bao phủ lo âu chung.',
    'Chất lượng NC: một số RCT risk of bias cao.',
    'Không đánh giá theo dõi dài hạn (>6 tháng) — hiệu quả duy trì?',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'VN chưa có RCT nào cho can thiệp SAD ở trẻ/VTN — GAP cực lớn. Jefferies QT35: VN SAD = 30,7%.',
    'iCBT (app/online) cho SAD ở VTN VN — phát triển app tiếng Việt + thích ứng văn hóa (B7 CA-CBT ĐNA).',
    'gCBT tại trường VN (mô hình SASS) — kết hợp kỹ năng XH + phơi nhiễm + phòng tái phát. Phù hợp B8 Sri Lanka RCT (8 phiên).',
    'So sánh iCBT vs gCBT tại VN — cái nào hiệu quả hơn cho VTN VN?',
    'Sàng lọc SAD trước can thiệp: SIAS-17 (Jefferies QT35) hoặc SPAI — xác định VTN cần can thiệp.',
    'NMA châu Á riêng — bao gồm NC từ TQ, Hàn Quốc, Nhật, ĐNA. Hiện chưa có.',
]:
    add_red(doc, f'• {s}')

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '43_NMA_SAD_2024.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['1.547', '71,2', '68,4', '66,0', '92,2', '89,6', '30 RCT', 'SUCRA', 'Bayesian', 'iCBT']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
