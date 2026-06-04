"""Build doc tong hop cac kien nghi cua phien 04/05/2026."""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/KIEN_NGHI_phien_04052026.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x00, 0x70, 0x30)
BLACK = RGBColor(0x00, 0x00, 0x00)
ORANGE = RGBColor(0xE6, 0x51, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para(text, color=BLACK, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = color
    r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def bullet(text, color=BLACK, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = color; r.font.size = Pt(12); r.italic = italic
    return p

# =====================================================
H('Kiến nghị tổng hợp — Phiên 04/05/2026', level=1)
H('— những đề xuất hành động cho thầy & cho dự án Lo âu HS THCS/THPT —', level=2)

para(
    'Doc này tổng hợp toàn bộ KIẾN NGHỊ em đã nêu trong phiên 04/05/2026 — '
    'bao gồm verify pending, sửa nguồn trích dẫn sai, format chuẩn, schedule '
    'agent tự động. Mỗi kiến nghị có ưu tiên 🔴 cao / 🟠 trung bình / 🟢 thấp.',
    italic=True
)

# =====================================================
H('A. Kiến nghị về VERIFY (chưa hoàn thành — pending)', level=2, color=RED)

para('A1. 🔴 CAO — Verify "Reckless" trong VN014 Hoàng Trung Học Bảng 3', bold=True)
bullet('Em chưa đọc trực tiếp PDF gốc bài Hoàng Trung Học 2025. Claim "Reckless = không lo âu" hiện chỉ dựa vào điểm thầy đã nêu — em chưa verify từ source.')
bullet('Hành động: thầy mở PDF gốc, chụp Bảng 3, gửi em verify. Hoặc thầy tự đọc + xác nhận. Nếu đúng là lỗi dịch, đáng gửi erratum tới tạp chí AJPR.')

para('A2. 🔴 CAO — Verify cell counts Bảng 3 Wen 2020 (QT008)', bold=True)
bullet('OR = 11,579 với KTC 95% [4,164; 32,194] — KTC rất rộng (UL/LL = 7,74). Khả năng nguyên nhân: cell tham chiếu (áp lực rất thấp + lo âu nặng) có ít quan sát.')
bullet('Em chưa có bản dịch QT008 trong corpus, chỉ có summary. Hành động: thầy đọc Bảng 3 PDF gốc Wen 2020, ghi rõ n từng ô — nếu cell tham chiếu < 10 thì OR phình to là tất yếu.')

para('A3. 🟠 TRUNG BÌNH — Verify con số 0,20 KTC 90% [0,12; 0,28] trong QT021 Brunborg', bold=True)
bullet('Em không tìm thấy nguyên văn này trong bản dịch QT021 hiện có. Có thể nằm trong supplementary material của bài gốc.')
bullet('Hành động: thầy gửi đoạn nguyên văn (số trang/bảng/section) để em verify; hoặc em sẽ đọc full bài gốc tiếng Anh nếu thầy cho phép tải bổ sung.')

para('A4. 🟠 TRUNG BÌNH — Verify tỷ lệ stress 65,5% trong VN014', bold=True)
bullet('Em viết "tỷ lệ 41,5–65,5%" theo claim của thầy nhưng chưa verify riêng số 65,5% trong tóm tắt VN014.')
bullet('Hành động: thầy xác nhận hoặc gửi đoạn trích.')

# =====================================================
H('B. Kiến nghị về SỬA NGUỒN TRÍCH DẪN (đã phát hiện sai)', level=2, color=RED)

para('B1. 🔴 CAO — Anderson 2025 (QT014) trích Daly 2022 SAI NGUỒN cho 31,9% RLLA', bold=True)
bullet('Anderson 2025 viết "31,9% VTN 13-18 có RLLA (Daly, 2022)". Nhưng Daly 2022 = "Prevalence of depression among adolescents in the US from 2009 to 2019" — Daly nghiên cứu TRẦM CẢM, không phải RLLA.')
bullet('Nguồn GỐC chính xác của 31,9% RLLA = Merikangas et al. (2010) NCS-A. Em đã thêm vào AUDIT phụ lục.')
bullet('Hành động: khi viết báo cáo CTH, KHÔNG dẫn Daly 2022 cho con số 31,9% RLLA — phải dẫn Merikangas et al. 2010 trực tiếp. Trích đầy đủ APA 7:')
bullet('"Merikangas, K. R., He, J. P., Burstein, M., Swanson, S. A., Avenevoli, S., Cui, L., Benjet, C., Georgiades, K., & Swendsen, J. (2010). Lifetime prevalence of mental disorders in U.S. adolescents: Results from the National Comorbidity Survey Replication–Adolescent Supplement (NCS-A). Journal of the American Academy of Child & Adolescent Psychiatry, 49(10), 980–989. https://doi.org/10.1016/j.jaac.2010.05.017"', italic=True)

para('B2. 🟠 TRUNG BÌNH — Soft thông tin về Ohannessian/Herres affiliation', bold=True)
bullet('Trong doc Tac_gia_khao_sat_coping_2007_Herres_Ohannessian_2015.docx, em viết Ohannessian là Professor tại UConn Health và Herres tại TCNJ — theo memory chuyên môn em, chưa verify từ web search.')
bullet('Hành động: thầy verify qua https://health.uconn.edu hoặc https://tcnj.edu trước khi trích. Em sẽ soft language trong bản update.')

para('B3. 🟠 TRUNG BÌNH — Quy tắc UL/LL > 4 = "KTC rộng"', bold=True)
bullet('Trong doc OR_11579_KTC, em viết bảng phân loại UL/LL theo quy ước thực hành. Quy ước này KHÔNG có cite chính thức từ một paper duy nhất — em đã hard-coded.')
bullet('Hành động: thầy có thể tham khảo Altman & Bland (1995, BMJ 311:485) về diễn giải KTC, nhưng phân loại UL/LL không phải chuẩn cố định.')

# =====================================================
H('C. Kiến nghị về QUY TRÌNH KIỂM TRA (theo memory)', level=2, color=NAVY)

para(
    'Theo feedback_verify_numbers_from_source.md (12/04/2026) và '
    'feedback_qa_pipeline_v2.md (pipeline 7 bước), em đề xuất:'
)

para('C1. 🔴 CAO — Mỗi khi em viết doc trả lời có số liệu, ÁP DỤNG TỐI THIỂU 4 BƯỚC:', bold=True)
bullet('Bước 2 (RAG): grep corpus tìm số.')
bullet('Bước 3 (verify_numbers): trace ngược về Tom-tat-tung-bai/{ID}_*.docx hoặc 03_Ban-dich/{ID}_*.docx.')
bullet('Bước 4 (Human reviewer flag): nếu không tìm thấy nguồn → ghi rõ "chưa verify" thay vì bịa.')
bullet('Bước 4b (Cross-ref audit): nếu cite paper khác → verify TRỰC TIẾP từ tóm tắt, không qua trí nhớ.')

para('C2. 🟠 TRUNG BÌNH — Chạy verify_numbers.py định kỳ', bold=True)
bullet('Script tham khảo: 06_Scripts/verify_vn001_vn029.py (đã có từ session VN022 13/04/2026).')
bullet('Em đề xuất build phiên bản generic: verify_numbers_v2.py — input = doc bất kỳ trong 01_Bao-cao/, output = list số liệu + nguồn match được.')

para('C3. 🟢 THẤP — Bổ sung canonical cho Pham 2024 (Huế)', bold=True)
bullet('Pham 2024 hiện em viết theo memory, chưa có canonical chính thức trong DB. Đề xuất canonical hóa làm VN029 hoặc VN030 nếu có PDF.')

# =====================================================
H('D. Kiến nghị về CẤU TRÚC TÀI LIỆU THAM KHẢO', level=2, color=NAVY)

para('D1. 🟠 TRUNG BÌNH — File DANH MỤC TLTK đã sort, cần thầy review', bold=True)
bullet('File mới: 01_Bao-cao/01_Gửi H_DANH MỤC TÀI LIỆU THAM KHẢO_DA_SAP_XEP.docx (55 KB).')
bullet('28 entries Tiếng Việt + 205 entries Tiếng Anh, sort alphabetically theo tác giả đầu tiên, giữ nguyên màu chữ (auto/đen/xanh).')
bullet('Hành động: thầy review thứ tự + đánh số 1, 2, 3... mới theo quy chuẩn của trường.')
bullet('Lưu ý nhỏ: 1 entry "AH_Geir Scott Brunborg..." có prefix "AH_" trong text nguồn → sort lên đầu EN list. Thầy chỉ cần delete prefix là đúng vị trí (Brunborg → B).')

para('D2. 🟠 TRUNG BÌNH — Format APA 7 cho 16 references chính', bold=True)
bullet('Em đã chuẩn bị danh mục APA 7 đầy đủ DOI/PMC trong AUDIT_4_doc_04052026 phần C — thầy có thể copy thẳng vào phần "Tiếng Anh" của DANH MỤC.')
bullet('Format chuẩn: Last, F. M., Last, F. M., & Last, F. M. (Year). Title in sentence case. Journal Name, Volume(Issue), pages. https://doi.org/...')

para('D3. 🟢 THẤP — Phân biệt rõ TLTK theo loại nguồn', bold=True)
bullet('Trong báo cáo CTH, có thể chia thêm: Tiếng Việt (sách/luận văn/tạp chí trong nước), Tiếng Anh (Q1/Q2/Q3/Q4), Báo cáo tổ chức (UNICEF, WHO, MOET, GBD), Trang web/Khác.')
bullet('Hiện tại file mới chỉ có 2 phần Việt/Anh — đủ cho luận văn/đề tài cấp cơ sở; nếu lên cấp trên cần phân nhỏ hơn.')

# =====================================================
H('E. Kiến nghị về AUTOMATION (cho session sau)', level=2, color=NAVY)

para('E1. 🟠 TRUNG BÌNH — /schedule agent kiểm tra pending verify items', bold=True)
bullet('Sau 1 tuần (11/05/2026), tự động ping em check 4 điểm pending verify (A1-A4 ở trên).')
bullet('Cách: thầy gõ /schedule trong session, em cấu hình cron weekly Monday 9:00.')

para('E2. 🟢 THẤP — Build mini-dashboard verify status cho từng bài corpus', bold=True)
bullet('Trong faceted browser tab "Sắp xếp đa diện" (đã build hôm 02/05), thêm 1 facet "Verify status" — Verified / Pending / No source.')
bullet('Hữu ích khi corpus có 91 bài và thầy cần biết bài nào đã verify, bài nào chưa.')

# =====================================================
H('F. Kiến nghị về PHIÊN LÀM VIỆC TIẾP THEO', level=2, color=NAVY)

para('F1. 🔴 CAO — Chuẩn hóa workflow trả lời câu hỏi', bold=True)
bullet('Format mới (đã thiết lập từ 23/04 + bổ sung hôm nay):')
bullet('  (1) Câu hỏi tô xanh + sửa typo;')
bullet('  (2) Background giải thích đen;')
bullet('  (3) CÂU TRẢ LỜI gom 1 chỗ tô xanh trước phụ lục;')
bullet('  (4) Phụ lục references APA 7 đầy đủ DOI/PMC;')
bullet('  (5) Self-audit cuối doc: ✓ verified / ⚠ pending / ❌ unverified;')
bullet('  (6) KHÔNG meta-note cuối doc.')

para('F2. 🟠 TRUNG BÌNH — Khi thầy hỏi về số liệu cụ thể, em LUÔN trace nguồn TRƯỚC', bold=True)
bullet('Theo feedback_verify_numbers_from_source.md: KHÔNG bao giờ "theo memory". Mỗi câu hỏi có số liệu → grep corpus trước, nếu không có trong corpus → ghi rõ "ngoài corpus, cần verify từ nguồn khác".')

para('F3. 🟢 THẤP — Memory update sau mỗi 3-5 doc trả lời', bold=True)
bullet('Hôm nay đã build 5 doc (4 trả lời + 1 AUDIT + 1 sort TLTK + 1 KIẾN NGHỊ). Em sẽ ghi vào project_session_summary_04052026.md sau khi hoàn tất.')

# =====================================================
H('G. Tóm tắt kiến nghị theo ưu tiên', level=2, color=BLUE)

para('🔴 CAO (làm ngay):', bold=True, color=RED)
bullet('A1. Verify "Reckless" trong VN014 Bảng 3.', color=RED)
bullet('A2. Verify cell counts Bảng 3 Wen 2020.', color=RED)
bullet('B1. KHI VIẾT, dẫn Merikangas 2010 thay Daly 2022 cho 31,9% RLLA.', color=RED)
bullet('C1. Áp dụng pipeline 4-bước verify cho mọi doc trả lời tiếp theo.', color=RED)
bullet('F1. Workflow chuẩn cho trả lời câu hỏi.', color=RED)

para('🟠 TRUNG BÌNH (làm trong tuần):', bold=True, color=ORANGE)
bullet('A3-A4. Verify QT021 KTC 90% + VN014 65,5%.', color=ORANGE)
bullet('B2-B3. Soft Ohannessian/Herres affiliation + UL/LL classification.', color=ORANGE)
bullet('C2. Build verify_numbers_v2.py generic.', color=ORANGE)
bullet('D1-D2. Review DANH MỤC TLTK đã sort + chèn 16 refs APA 7.', color=ORANGE)
bullet('E1. /schedule agent weekly verify check.', color=ORANGE)
bullet('F2. Trace nguồn trước khi viết.', color=ORANGE)

para('🟢 THẤP (khi rảnh):', bold=True, color=GREEN)
bullet('C3. Canonical hóa Pham 2024 (Huế).', color=GREEN)
bullet('D3. Phân nhỏ TLTK theo loại nguồn.', color=GREEN)
bullet('E2. Verify status facet trong faceted browser.', color=GREEN)
bullet('F3. Memory update.', color=GREEN)

# =====================================================
H('H. Tham chiếu memory đã đọc', level=2)
para('Em đã đọc các memory sau (tại C:\\Users\\HLC\\.claude\\projects\\c--Users-HLC-OneDrive-read-books-Lo-au\\memory\\):', italic=True, size=11)
bullet('feedback_verify_numbers_from_source.md (12/04/2026) — 6 fabricated claims phát hiện từ v5.', italic=True)
bullet('feedback_qa_pipeline_v2.md — Pipeline 7 bước QA validated.', italic=True)
bullet('feedback_doc_format_blue_answer.md (23/04) — Format CÂU TRẢ LỜI tô xanh trước phụ lục.', italic=True)
bullet('feedback_translation_principles_v2.md — 10 nguyên tắc dịch (đặc biệt #9 cross-ref audit).', italic=True)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
