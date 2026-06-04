# -*- coding: utf-8 -*-
"""QT26 UK NHS + QT27 Nature — chuẩn CTH v5, 2-3 bảng/bài"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
def make_doc():
    d = Document(); s = d.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in d.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return d
def rb(d,t):
    p=d.add_paragraph();r=p.add_run(t);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=RED
def bl(d,t,bold=False):
    p=d.add_paragraph();r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=BLUE;r.bold=bold
def rh2(d,t):
    h=d.add_heading(t,level=2)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=RED
def bh3(d,t):
    h=d.add_heading(t,level=3)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=BLUE
def shade(c,color):
    s=OxmlElement('w:shd');s.set(qn('w:fill'),color);s.set(qn('w:val'),'clear');c._tc.get_or_add_tcPr().append(s)
def set_w(c,w):
    tcW=c._tc.get_or_add_tcPr();we=OxmlElement('w:tcW');we.set(qn('w:w'),str(int(w*567)));we.set(qn('w:type'),'dxa');tcW.append(we)
def tbl(d,headers,rows,widths):
    t=d.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):set_w(row.cells[ci],widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci];c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:r.font.name='Times New Roman';r.font.size=Pt(10)

# ===== QT26: UK NHS =====
d=make_doc()
bl(d,'Tóm tắt bài QT-26',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Báo cáo \u00ab Thống kê sức khỏe tâm thần: Tỷ lệ, dịch vụ và tài trợ tại Anh \u00bb (Mental Health Statistics: Prevalence, Services and Funding in England), do Carl Baker, House of Commons Library, UK Parliament (2025). 46 trang. Nguồn chính thức Quốc hội Anh, dựa trên NHS Mental Health Survey 2025 và nhiều khảo sát quốc gia.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Báo cáo tổng hợp chính sách (policy briefing) từ Quốc hội Anh, sử dụng dữ liệu NHS Digital, Health Survey for England, và nhiều nguồn chính thức. Nói cách khác, đây không phải nghiên cứu gốc mà là tổng hợp dữ liệu quốc gia phục vụ hoạch định chính sách \u2014 nguồn đáng tin cậy nhất về SKTT tại Anh.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1. Xu hướng SKTT thanh niên Anh (NHS Survey):', bold=True)
tbl(d,['Chỉ số','2014','2025','Thay đổi'],
    [['Rối loạn TT phổ biến (16\u201324 tuổi)','18,9%','25,8%','+36% (11 năm)'],
     ['Nữ 16\u201324','','36,1%','Gấp 2,2 lần nam'],
     ['Nam 16\u201324','','16,3%',''],
     ['VTN 17\u201319 đủ tiêu chuẩn rối loạn TT','','~25%','1/4 VTN'],
     ['Tự hại suốt đời (mọi tuổi)','','10,3%','Tăng gấp 4 từ 2000'],
     ['Nữ 16\u201324 tự hại','','31,7%','Rất cao']],
    widths=[5.0,2.0,2.0,3.0])
d.add_paragraph()

bl(d,'Bảng 2. So sánh xu hướng Anh với các nước:', bold=True)
tbl(d,['Quốc gia','Giai đoạn','Xu hướng lo âu/SKTT','Nguồn'],
    [['Anh (NHS)','2014\u20132025','+36% rối loạn TT 16\u201324','Bài này'],
     ['Hoa Kỳ (JAACAP)','2013\u20132021','Lo âu +100% (AOR=2,17)','QT23'],
     ['Hoa Kỳ (NSCH)','2016\u20132023','Lo âu +61%','Bài 04'],
     ['Na Uy','2011\u20132024','Tăng liên tục 13 năm','QT21'],
     ['Ireland','2012\u20132019','Tăng, nữ nhanh hơn','QT32'],
     ['Hàn Quốc','2006\u20132022','Giảm trước, tăng sau COVID','QT34'],
     ['Việt Nam (HTH)','2021 vs 2023','41,5% \u2192 25,4% (COVID giảm)','VN14']],
    widths=[3.0,2.5,4.0,2.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : Xu hướng tăng SKTT ở VTN Anh (+36% trong 11 năm) nhất quán với Mỹ, Na Uy, Ireland. Nữ bị ảnh hưởng nặng hơn ở TẤT CẢ các nước. Đáng chú ý, tự hại ở nữ 16\u201324 Anh (31,7%) rất cao \u2014 so sánh Danh Lâm 2022 VN (tự làm đau 10%, cố tự tử 1,4%).')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*Xu hướng toàn cầu nhất quán.* Tất cả các nước phương Tây đều báo cáo tăng SKTT ở VTN trong 10\u201315 năm qua. Anh không phải ngoại lệ.')
bl(d,'*Nữ \u2014 nhóm đặc biệt dễ tổn thương.* 36,1% nữ 16\u201324 mắc rối loạn TT (gấp 2,2 lần nam) \u2014 cần can thiệp nhạy cảm giới.')
bl(d,'*Tự hại tăng gấp 4 từ 2000.* Xu hướng đáng báo động, phù hợp với dữ liệu tự hại ở VN (Danh Lâm: 10%).')

rb(d,'Kết luận')
bl(d,'Dữ liệu NHS Anh cho thấy 25,8% thanh niên 16\u201324 mắc rối loạn tâm thần phổ biến \u2014 tăng 36% trong 11 năm \u2014 với nữ bị ảnh hưởng gấp đôi nam và tự hại tăng gấp 4. Xu hướng nhất quán với Mỹ, Na Uy, Ireland, gợi ý cuộc khủng hoảng SKTT VTN là toàn cầu.',bold=True)

rh2(d,'Phản biện')
bl(d,'Nguồn chính thức Quốc hội Anh \u2014 đáng tin cậy nhất. Dữ liệu NHS quốc gia. Tuy nhiên, báo cáo chính sách, không phải NC peer-reviewed. Chỉ Anh \u2014 bối cảnh khác VN.')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'So sánh hệ thống NHS Anh với VN về tiếp cận dịch vụ SKTT. Phân tích xu hướng tự hại ở VTN VN.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50 Trung bình-Khá. Nguồn chính thức nhưng không peer-reviewed.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT26_UK_NHS_2025.docx')
sys.stderr.write('QT26 OK\n')

# ===== QT27: NATURE SOCIAL MEDIA =====
d=make_doc()
bl(d,'Tóm tắt bài QT-27',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Sử dụng mạng xã hội ở thanh thiếu niên có và không có tình trạng sức khỏe tâm thần \u00bb (Social Media Use in Adolescents with and without Mental Health Conditions), xuất bản trên Nature Human Behaviour, tập 9, tháng 6/2025, trang 1283\u20131299 (Q1, IF \u2248 24,0). N = 3.340 VTN (11\u201319 tuổi), mẫu đại diện quốc gia Vương quốc Anh. Bao gồm đánh giá chẩn đoán bởi chuyên gia lâm sàng. 21 trang.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Nghiên cứu cắt ngang mẫu đại diện quốc gia UK, N = 3.340 VTN 11\u201319 tuổi. Điểm mạnh đặc biệt: SKTT được đánh giá bằng CHẨN ĐOÁN LÂM SÀNG (clinical raters) \u2014 không chỉ sàng lọc tự báo cáo. Nói cách khác, đây là một trong số ít NC phân biệt rõ VTN "có rối loạn SKTT được chẩn đoán" vs "không có" khi đánh giá hành vi mạng xã hội.',bold=True)

bl(d,'Tổng quan tài liệu cho thấy mối quan hệ mạng xã hội \u2192 SKTT đã được báo cáo rộng rãi nhưng đa số từ dữ liệu tự báo cáo, thiếu chẩn đoán lâm sàng. Nature Human Behaviour (IF = 24) là tạp chí có impact factor cao nhất trong tất cả bài NC của Đề tài.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1. So sánh hành vi mạng xã hội theo tình trạng SKTT:', bold=True)
tbl(d,['Nhóm VTN','Thời gian dùng MXH','Mức hài lòng','Số bạn online'],
    [['Có rối loạn SKTT (chẩn đoán)','Nhiều hơn đáng kể','Ít hài lòng hơn','Nhiều hơn nhưng kém hài lòng'],
     ['Không rối loạn SKTT','Ít hơn','Hài lòng hơn','Ít hơn nhưng hài lòng hơn'],
     ['48% rối loạn SKTT','Triệu chứng trước 18 tuổi','','']],
    widths=[4.0,3.5,2.5,3.5])
d.add_paragraph()

bl(d,'Bảng 2. So sánh bằng chứng về mạng xã hội/screen time từ nhiều NC:', bold=True)
tbl(d,['Nguồn','Thiết kế','Phát hiện chính','IF'],
    [['Nature 2025 (bài này)','Cắt ngang + chẩn đoán lâm sàng','VTN có RLTT dùng MXH nhiều hơn','24,0'],
     ['BJCP 2025 (QT22)','DỌC','Screen time DỰ BÁO lo âu/trầm cảm sau 1 năm','3,0'],
     ['JAMA 2024 (QT33)','RCT phân tích thứ cấp','Giảm screen time cải thiện SKTT','13,0'],
     ['Chen 2023 (bài 07)','Cắt ngang, n=63.205','Rối loạn game OR = 5,00','4,4'],
     ['Norway 2025 (QT21)','Xu hướng 13 năm','MXH giải thích xu hướng tăng','5,4'],
     ['Hoàng Trung Học VN','Cắt ngang, n=8.473','Điện tử Beta = 0,176','']],
    widths=[3.5,3.5,4.0,1.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : Sáu nguồn bằng chứng hội tụ: (1) Nature 2025 \u2014 VTN có RLTT dùng MXH nhiều hơn (chẩn đoán lâm sàng); (2) BJCP 2025 \u2014 screen time DỰ BÁO lo âu (dọc); (3) JAMA 2024 \u2014 giảm screen time cải thiện SKTT (RCT); (4) Chen 2023 \u2014 game OR = 5,00; (5) Norway 2025 \u2014 MXH giải thích xu hướng tăng; (6) Hoàng Trung Học VN \u2014 điện tử Beta = 0,176. Nói cách khác, bằng chứng từ nhiều thiết kế (cắt ngang, dọc, RCT) và nhiều nước đều chỉ cùng một hướng: mạng xã hội/screen time ảnh hưởng tiêu cực đến SKTT VTN.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*Chẩn đoán lâm sàng \u2014 điểm mạnh vượt trội.* Đa số NC chỉ dùng sàng lọc tự báo cáo. Nature 2025 dùng chuyên gia lâm sàng đánh giá \u2014 kết quả đáng tin cậy hơn.')
bl(d,'*48% rối loạn SKTT bắt đầu trước 18 tuổi.* Nhấn mạnh tầm quan trọng can thiệp sớm ở VTN.')
bl(d,'*Nghịch lý: nhiều bạn online nhưng kém hài lòng.* VTN có RLTT có nhiều "bạn" online nhưng kém hài lòng với các mối quan hệ \u2014 gợi ý chất lượng kết nối quan trọng hơn số lượng.')

rb(d,'Kết luận')
bl(d,'Dữ liệu từ Nature Human Behaviour (IF = 24) trên 3.340 VTN UK với chẩn đoán lâm sàng, cho thấy VTN có rối loạn SKTT dùng mạng xã hội nhiều hơn và kém hài lòng hơn, gợi ý rằng can thiệp quản lý mạng xã hội cần nhắm đến cả thời gian lẫn chất lượng tương tác \u2014 đặc biệt cho VTN đã có dấu hiệu rối loạn.',bold=True)

rh2(d,'Phản biện')
bl(d,'Nature Q1 IF = 24 \u2014 tạp chí TOP toàn cầu. n = 3.340 quốc gia UK. Chẩn đoán lâm sàng \u2014 vượt trội so với tự báo cáo. Tuy nhiên, cắt ngang \u2014 không suy luận nhân quả (MXH gây RLTT hay RLTT dẫn đến dùng MXH nhiều hơn?). Chỉ UK \u2014 văn hóa MXH có thể khác châu Á.')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'NC tương tự tại VN: so sánh VTN có/không RLTT về hành vi MXH. Kết hợp chẩn đoán lâm sàng + dữ liệu sử dụng MXH thực tế (không chỉ tự báo cáo). Can thiệp quản lý chất lượng tương tác MXH.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. Nature Q1 IF = 24, mẫu quốc gia, chẩn đoán lâm sàng.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT27_Nature_SocialMedia_2025.docx')
sys.stderr.write('QT27 OK\n')
