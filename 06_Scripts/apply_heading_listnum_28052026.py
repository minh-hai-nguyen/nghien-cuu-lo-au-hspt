# -*- coding: utf-8 -*-
"""Apply Heading 2 to auto-numbered paragraphs trong MỞ ĐẦU + KẾT LUẬN
ma chua co Heading style. Sau do rebuild TOC.
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')


def in_table(p):
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag in (qn('w:tbl'), qn('w:tc')):
            return True
        parent = parent.getparent()
    return False


def has_numbering(p):
    """Check if paragraph has Word auto-numbering (w:numPr)."""
    pPr = p._element.find(qn('w:pPr'))
    if pPr is None: return False
    numPr = pPr.find(qn('w:numPr'))
    return numPr is not None


d = Document(FILE)

# Find MỞ ĐẦU and KẾT LUẬN positions
mo_dau_idx = None
chuong1_idx = None
ket_luan_idx = None
phu_luc_idx = None
for i, p in enumerate(d.paragraphs):
    txt = p.text.strip().upper()
    if txt == 'MỞ ĐẦU' and mo_dau_idx is None:
        mo_dau_idx = i
    elif re.match(r'^CHƯƠNG\s+1\s*$', txt) and chuong1_idx is None and i > (mo_dau_idx or 0):
        chuong1_idx = i
    elif txt == 'KẾT LUẬN VÀ KIẾN NGHỊ' and ket_luan_idx is None:
        ket_luan_idx = i
    elif txt == 'TÀI LIỆU THAM KHẢO' and phu_luc_idx is None:
        phu_luc_idx = i

print(f"MỞ ĐẦU: para {mo_dau_idx}")
print(f"CHƯƠNG 1: para {chuong1_idx}")
print(f"KẾT LUẬN: para {ket_luan_idx}")
print(f"TLTK: para {phu_luc_idx}")


# Apply Heading 2 to auto-numbered paragraphs in MỞ ĐẦU and KẾT LUẬN sections
applied = []
sections = [
    ('MỞ ĐẦU', mo_dau_idx + 1 if mo_dau_idx else 100, chuong1_idx if chuong1_idx else 220),
    ('KẾT LUẬN', ket_luan_idx + 1 if ket_luan_idx else 1140, phu_luc_idx if phu_luc_idx else 1199),
]

for sect_name, start, end in sections:
    print(f"\n=== Section {sect_name} (para {start}-{end}) ===")
    for i in range(start, min(end, len(d.paragraphs))):
        p = d.paragraphs[i]
        if in_table(p): continue
        text = p.text.strip()
        if not text or len(text) > 250: continue
        style = p.style.name if p.style else 'Normal'

        # Already a heading - skip
        if 'Heading' in style or 'Tiêu đề' in style:
            continue

        # Check for auto-numbering OR text that looks like heading
        is_numbered = has_numbering(p)
        # Or text starts with no number but is a heading-like keyword
        # (already handled by apply_heading_v2)

        if is_numbered:
            # Likely a heading - check if it has bold runs (heading-like)
            bold_runs = sum(1 for r in p.runs if r.text.strip() and r.bold)
            total_runs = sum(1 for r in p.runs if r.text.strip())
            if total_runs == 0: continue
            is_bold = bold_runs >= max(1, total_runs - 1)
            # Apply Heading 2 (most common in MỞ ĐẦU/KẾT LUẬN)
            try:
                p.style = d.styles['Heading 2']
                applied.append((i, text[:80], style, 'auto-numbered'))
            except: pass

print(f"\nApplied {len(applied)} Heading 2 to auto-numbered paragraphs:")
for idx, txt, old_style, reason in applied:
    print(f"  Para {idx} ({reason}, was {old_style}): {txt}")

d.save(FILE)
print(f"\nSaved: {FILE}")
