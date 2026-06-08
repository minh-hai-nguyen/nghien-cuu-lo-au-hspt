# -*- coding: utf-8 -*-
"""Outline IMRaD bai Q4 — Latent Profile Analysis (LPA).
Song ngu Anh-Viet, ~3000 tu, reuse Q2 v1 + add LPA-specific methodology."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'Outline_Q4_IMRaD_SongNgu_v1_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def TITLE(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def EN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def VN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def PB(en, vn):
    EN(en); VN(vn)

def BB(en, vn, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('• ' + en); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('▸ ' + vn); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def NOTE(en, vn):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('[NOTE TO NCS + co-authors] ' + en)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('[GHI CHÚ cho NCS + đồng tác giả] ' + vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)


# =====================================================
TITLE('OUTLINE IMRaD — BÀI Q4')
TITLE('Anxiety phenotype profiles in Vietnamese lower secondary '
      'students: A latent profile analysis with risk-protective '
      'indicator integration')
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Hồ sơ phân loại lo âu ở học sinh trung học cơ sở Việt '
              'Nam: Phân tích hồ sơ tiềm ẩn tích hợp các chỉ báo '
              'nguy cơ và bảo vệ')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Outline song ngữ Anh-Việt v1 — soạn 08/06/2026 — '
              'kế hoạch năm 2027')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

NOTE('This is an IMRaD outline (structure + key claims, not full prose). '
     'Authors: Cong TH (1st), [Nguyen MD], Hai-Nguyen Minh. Target journal: '
     'Journal of Psychopathology and Behavioral Assessment OR BMC '
     'Psychiatry. Q4 is planned for 2027 (after Q2 + Q3 published), '
     'pending data sufficiency. Word target full draft: 7,000-8,500.',
     'Đây là outline IMRaD (cấu trúc + luận điểm chính, không phải bản '
     'thảo đầy đủ). Tác giả: Công TH (chủ trì), [Nguyễn MD], Hải-Nguyễn '
     'Minh. Tạp chí đích: Journal of Psychopathology and Behavioral '
     'Assessment HOẶC BMC Psychiatry. Q4 dự kiến năm 2027 (sau khi Q2 '
     '+ Q3 đã công bố), tùy theo độ đủ dữ liệu. Word target draft đầy '
     'đủ: 7.000-8.500 từ.')


# =====================================================
H1('ABSTRACT (300 words target)')

H2('Background')
PB('Variable-centered analyses (SEM in our Q2 and Q3) test how '
   'risk and protective factors RELATE to anxiety outcomes, but '
   'cannot detect whether students cluster into distinct '
   '"phenotype profiles" — qualitatively different patterns of '
   'risk-protective configurations. Person-centered methods such '
   'as Latent Profile Analysis (LPA) complement variable-centered '
   'approaches by identifying unobserved subgroups within the '
   'population. To date, LPA-based phenotyping of adolescent '
   'anxiety has not been applied to Vietnamese cohorts.',
   'Các phân tích lấy biến làm trung tâm (SEM trong Q2 và Q3 của '
   'chúng tôi) kiểm chứng cách các yếu tố nguy cơ và bảo vệ LIÊN '
   'QUAN tới các kết quả lo âu, nhưng không phát hiện được liệu '
   'học sinh có gộp thành các "hồ sơ phân loại" riêng biệt hay '
   'không — các mẫu hình cấu hình nguy cơ-bảo vệ khác nhau về '
   'mặt định tính. Các phương pháp lấy con người làm trung tâm '
   'như Phân tích Hồ sơ Tiềm ẩn (LPA) bổ sung cho cách tiếp cận '
   'lấy biến làm trung tâm bằng cách nhận diện các nhóm con '
   'không quan sát được trong dân số. Cho đến nay, phân loại '
   'hồ sơ lo âu vị thành niên dựa trên LPA chưa được áp dụng '
   'cho cohort Việt Nam.')

H2('Methods')
PB('Cross-sectional survey of 1,352 Vietnamese lower secondary '
   'students (same dataset as Q2 and Q3). Ten indicator variables '
   '(three risk factors, four protective factors, three DSM-5 '
   'anxiety subtype scores) were submitted to LPA. Models with '
   'K = 2, 3, 4, 5, 6 profiles were compared using BIC, AIC, '
   'entropy, and Lo-Mendell-Rubin Likelihood Ratio Test (LMR-'
   'LRT). The optimal solution was interpreted substantively and '
   'externally validated against gender, grade, and school '
   'distributions.',
   'Khảo sát cắt ngang 1.352 học sinh THCS Việt Nam (cùng bộ dữ '
   'liệu với Q2 và Q3). Mười biến chỉ báo (ba yếu tố nguy cơ, '
   'bốn yếu tố bảo vệ, ba điểm phân loại lo âu DSM-5) được đưa '
   'vào LPA. Mô hình K = 2, 3, 4, 5, 6 hồ sơ được so sánh sử '
   'dụng BIC, AIC, entropy, và Kiểm định Lo-Mendell-Rubin '
   '(LMR-LRT). Giải pháp tối ưu được diễn giải về mặt thực '
   'chất và validated bên ngoài qua phân bố giới, khối lớp, và '
   'trường học.')

H2('Results (planned)')
PB('We anticipate identifying 4-5 distinct anxiety phenotype '
   'profiles — for example: (1) low-risk + high-protective + '
   'low-anxiety (resilient majority), (2) high-academic-stress + '
   'low-self-esteem + moderate-anxiety (achievement-driven), (3) '
   'high-bullying + low-school-membership + high-SAD (peer-'
   'victimization profile), (4) high-smartphone + low-parental-'
   'support + moderate-SocAD (digital-isolation), and possibly a '
   '(5) high-on-all-risk + low-on-all-protective + severe-anxiety '
   'profile. Profile distributions will be tested for gender, '
   'grade, and school differences.',
   'Chúng tôi kỳ vọng nhận diện được 4-5 hồ sơ phân loại lo âu '
   'riêng biệt — ví dụ: (1) thấp nguy cơ + cao bảo vệ + thấp lo '
   'âu (đa số kiên cường), (2) cao áp lực học tập + thấp tự '
   'trọng + vừa lo âu (theo đuổi thành tích), (3) cao bắt nạt + '
   'thấp gắn bó trường + cao SAD (nạn nhân hóa từ bạn bè), (4) '
   'cao điện thoại + thấp hỗ trợ cha mẹ + vừa SocAD (cô lập số), '
   'và có thể (5) cao tất cả nguy cơ + thấp tất cả bảo vệ + lo '
   'âu nặng. Phân bố hồ sơ sẽ được kiểm chứng theo giới, khối '
   'lớp, và trường.')

H2('Conclusions (planned)')
PB('Phenotype profiles will inform targeted prevention — different '
   'profiles likely benefit from different intervention components. '
   'Together with Q2 (integrated mechanisms) and Q3 (gender '
   'invariance), Q4 completes a "hat-trick" three-paper program '
   'characterizing Vietnamese adolescent anxiety from variable-, '
   'group-, and person-centered perspectives.',
   'Các hồ sơ phân loại sẽ cung cấp thông tin cho phòng ngừa có '
   'mục tiêu — các hồ sơ khác nhau có thể hưởng lợi từ các thành '
   'phần can thiệp khác nhau. Cùng với Q2 (cơ chế tích hợp) và '
   'Q3 (bất biến giới), Q4 hoàn thành chương trình ba bài "hat-'
   'trick" mô tả đặc điểm lo âu vị thành niên Việt Nam từ ba '
   'góc nhìn: biến, nhóm, và con người làm trung tâm.')


# =====================================================
H1('1. INTRODUCTION (target 1,000-1,200 words full draft)')

H2('1.1 — Variable-centered vs person-centered approaches')
PB('Most adolescent anxiety research uses variable-centered '
   'methods (correlation, regression, SEM) that estimate average '
   'relationships in the population. These methods assume the '
   'sample is homogeneous — a single set of pathway coefficients '
   'applies to all individuals. Person-centered methods (LPA, '
   'LCA, GMM) relax this assumption by allowing the population '
   'to contain unobserved subgroups with qualitatively different '
   'patterns. Both approaches answer complementary scientific '
   'questions (Bauer 2007; Lanza & Cooper 2016).',
   'Hầu hết nghiên cứu lo âu vị thành niên sử dụng các phương '
   'pháp lấy biến làm trung tâm (tương quan, hồi quy, SEM) ước '
   'lượng các quan hệ trung bình trong dân số. Các phương pháp '
   'này giả định mẫu đồng nhất — một tập hợp hệ số đường dẫn áp '
   'dụng cho tất cả cá nhân. Các phương pháp lấy con người làm '
   'trung tâm (LPA, LCA, GMM) nới lỏng giả định này bằng cách '
   'cho phép dân số chứa các nhóm con không quan sát được với '
   'các mẫu hình khác nhau về định tính. Cả hai cách tiếp cận '
   'trả lời các câu hỏi khoa học bổ sung lẫn nhau (Bauer 2007; '
   'Lanza & Cooper 2016).')

H2('1.2 — Why person-centered analyses matter for prevention')
PB('Anxiety disorders are common in adolescence — half of all '
   'lifetime cases begin by age 14 (Kessler et al. 2005), and '
   'rates are highest in girls (Hankin et al. 2007; McLean & '
   'Anderson 2009). Cross-national epidemiology shows the burden '
   'varies widely across countries (Jefferies & Ungar 2020). '
   'Public health practice often allocates intervention resources '
   'by risk profile — high-risk youth receive intensive programs, '
   'low-risk youth receive universal prevention. Identifying '
   'meaningful phenotype profiles helps match the intensity and '
   'content of intervention to student needs, avoiding both '
   'under-treatment of high-need students and over-treatment of '
   'low-need students.',
   'Rối loạn lo âu phổ biến ở vị thành niên — một nửa số ca '
   'khởi phát trước 14 tuổi (Kessler và cs. 2005), và tỷ lệ '
   'cao nhất ở nữ (Hankin và cs. 2007; McLean & Anderson 2009). '
   'Dịch tễ xuyên quốc gia cho thấy gánh nặng khác biệt rộng '
   'rãi giữa các quốc gia (Jefferies & Ungar 2020). Thực hành '
   'y tế công cộng thường phân bổ nguồn lực can thiệp theo hồ '
   'sơ nguy cơ — thanh thiếu niên nguy cơ cao nhận chương trình '
   'chuyên sâu, thanh thiếu niên nguy cơ thấp nhận phòng ngừa '
   'phổ quát. Nhận diện các hồ sơ phân loại có ý nghĩa giúp '
   'khớp cường độ và nội dung can thiệp với nhu cầu học sinh, '
   'tránh đồng thời việc thiếu can thiệp cho học sinh nhu cầu '
   'cao và can thiệp thừa cho học sinh nhu cầu thấp.')

H2('1.3 — Gap: no Vietnamese phenotype profiling')
PB('LPA-based phenotyping of adolescent anxiety has been '
   'conducted in samples from China, US, Australia, and several '
   'European countries — but not Vietnam. Given the distinctive '
   'risk-protective configuration of Vietnamese lower secondary '
   'students (Confucian academic culture + tam giao + rapid '
   'digital transformation), it is plausible that Vietnamese '
   'phenotype profiles differ from Western or other Asian '
   'profiles. This study fills that gap.',
   'Phân loại hồ sơ lo âu vị thành niên dựa trên LPA đã được '
   'thực hiện trên các mẫu từ Trung Quốc, Mỹ, Úc, và một số '
   'quốc gia châu Âu — nhưng chưa có Việt Nam. Với cấu hình '
   'nguy cơ-bảo vệ đặc trưng của học sinh THCS Việt Nam (văn '
   'hóa học thuật Nho giáo + tam giáo + chuyển đổi số nhanh), '
   'có thể các hồ sơ phân loại của Việt Nam khác biệt so với '
   'hồ sơ phương Tây hoặc châu Á khác. Nghiên cứu này lấp đầy '
   'khoảng trống đó.')

H2('1.4 — Research questions (exploratory)')
PB('Unlike Q2 and Q3 which test pre-registered hypotheses, Q4 '
   'asks three exploratory research questions. RQ1: How many '
   'distinct anxiety phenotype profiles are present in the '
   'sample? RQ2: How are profiles characterized in terms of risk '
   'and protective factor levels? RQ3: How do profile membership '
   'distributions vary by gender, grade, and school?',
   'Khác với Q2 và Q3 kiểm chứng các giả thuyết đăng ký trước, '
   'Q4 đặt ba câu hỏi nghiên cứu thám sát. RQ1: Có bao nhiêu hồ '
   'sơ phân loại lo âu khác biệt trong mẫu? RQ2: Các hồ sơ được '
   'đặc trưng như thế nào về mức độ các yếu tố nguy cơ và bảo '
   'vệ? RQ3: Phân bố thành viên hồ sơ thay đổi theo giới, khối '
   'lớp, và trường như thế nào?')

NOTE('LPA is exploratory by design — pre-registration is recommended '
     'but the focus is on confirming the EXPLORATORY procedure (LMR-'
     'LRT criteria) rather than specific profile predictions.',
     'LPA là thám sát theo thiết kế — đăng ký trước được khuyến '
     'nghị nhưng trọng tâm là xác nhận QUY TRÌNH thám sát (tiêu '
     'chí LMR-LRT) thay vì các dự đoán hồ sơ cụ thể.')


# =====================================================
H1('2. METHODS (target 1,000-1,200 words full draft)')

H2('2.1 — Participants')
PB('Reuse Q2 v1 participant description verbatim (1,352 students, '
   'Nhat Tan + Tay Mo, ages 11-14).',
   'Sử dụng lại phần Participants của Q2 v1 nguyên văn (1.352 '
   'học sinh, Nhật Tân + Tây Mỗ, 11-14 tuổi).')

H2('2.2 — Measures and indicator selection')
PB('Reuse Q2 v1 measures section. For LPA, select 10 indicators: '
   '3 risk (bullying OBVQ, academic stress ESSA, smartphone SAS-'
   'SV), 4 protective (school membership PSSM, parental support '
   'MSPSS, peer support MSPSS, self-esteem RSES), 3 anxiety '
   'subtype scores (GAD, SAD, SocAD from RCADS). All standardized '
   'to mean = 50, SD = 10 for LPA estimation comparability.',
   'Sử dụng lại phần Measures của Q2 v1. Cho LPA, chọn 10 chỉ '
   'báo: 3 nguy cơ (bắt nạt OBVQ, áp lực học tập ESSA, điện '
   'thoại SAS-SV), 4 bảo vệ (gắn bó trường PSSM, hỗ trợ cha mẹ '
   'MSPSS, hỗ trợ bạn bè MSPSS, tự trọng RSES), 3 điểm phân '
   'loại lo âu (GAD, SAD, SocAD từ RCADS). Tất cả chuẩn hóa '
   'thành trung bình = 50, SD = 10 để so sánh ước lượng LPA.')

H2('2.3 — LPA analytic strategy')

BB('Step 1: Indicator preparation. Confirm indicator normality '
   '(skewness/kurtosis within ±1), missing data check, and '
   'standardization to T-score scale.',
   'Bước 1: Chuẩn bị chỉ báo. Xác nhận tính chuẩn của chỉ báo '
   '(độ lệch/độ nhọn trong khoảng ±1), kiểm tra dữ liệu khuyết, '
   'và chuẩn hóa sang thang T-score.', 0)

BB('Step 2: Compare K-profile solutions. Estimate LPA with K = '
   '2, 3, 4, 5, 6 profiles. Use maximum likelihood with 500 '
   'random starts to avoid local maxima.',
   'Bước 2: So sánh các giải pháp K hồ sơ. Ước lượng LPA với '
   'K = 2, 3, 4, 5, 6 hồ sơ. Dùng maximum likelihood với 500 '
   'điểm khởi đầu ngẫu nhiên để tránh cực đại địa phương.', 0)

BB('Step 3: Model selection criteria following Nylund et al. '
   '(2007) recommendations. Compare BIC (lower = better), AIC '
   '(lower = better), entropy (higher = better; aim ≥ 0.80), '
   'LMR-LRT (significant = larger model preferred), and '
   'parsimony. No single criterion dictates — consider '
   'convergence across criteria + substantive interpretability.',
   'Bước 3: Tiêu chí chọn mô hình theo khuyến nghị của Nylund '
   'và cs. (2007). So sánh BIC (thấp hơn = tốt hơn), AIC '
   '(thấp hơn = tốt hơn), entropy (cao hơn = tốt hơn; mục tiêu '
   '≥ 0,80), LMR-LRT (có ý nghĩa = mô hình lớn hơn được ưu '
   'tiên), và sự tiết kiệm. Không tiêu chí nào quyết định một '
   'mình — cân nhắc sự hội tụ qua các tiêu chí + tính diễn '
   'giải thực chất.', 0)

BB('Step 4: Profile interpretation. For optimal K, examine '
   'standardized mean profile across indicators. Name each '
   'profile descriptively (e.g., "high-stress + low-support + '
   'moderate-anxiety"). Estimate profile prevalence (%) in '
   'sample.',
   'Bước 4: Diễn giải hồ sơ. Với K tối ưu, kiểm tra trung bình '
   'chuẩn hóa của hồ sơ qua các chỉ báo. Đặt tên mỗi hồ sơ '
   'theo mô tả (vd "cao áp lực + thấp hỗ trợ + vừa lo âu"). '
   'Ước lượng tỷ lệ hồ sơ (%) trong mẫu.', 0)

BB('Step 5: External validation. Test profile membership '
   'differences by gender (chi-square), grade (chi-square), '
   'and school (chi-square). Cohen\'s w effect sizes for each '
   'comparison.',
   'Bước 5: Validation bên ngoài. Kiểm chứng khác biệt thành '
   'viên hồ sơ theo giới (chi-square), khối lớp (chi-square), '
   'và trường (chi-square). Effect size Cohen\'s w cho mỗi so '
   'sánh.', 0)

H2('2.4 — Software / Phần mềm')
PB('Estimation in Mplus 8.10 (or R package "tidyLPA" v1.1+ for '
   'reproducibility). Bootstrap LMR-LRT with 500 replications.',
   'Ước lượng bằng Mplus 8.10 (hoặc gói R "tidyLPA" v1.1+ để '
   'có thể tái tạo). Bootstrap LMR-LRT với 500 lần lặp.')

H2('2.5 — Ethics')
PB('Reuse Q2 v1 ethics section. [TBD — Q3-6 BLOCKING: IRB '
   'number and date pending official letter from HNUE Ethics '
   'Committee.]',
   'Sử dụng lại phần Ethics của Q2 v1. [CHỜ XÁC NHẬN — Q3-6 '
   'BLOCKING: số quyết định và ngày của HĐ đạo đức HNUE chờ '
   'thư chính thức.]')


# =====================================================
H1('3. RESULTS (target 1,500-1,800 words full draft)')

H2('3.1 — Indicator descriptives')
PB('Means, SDs, skewness, kurtosis, and intercorrelations of '
   '10 indicators (Table 1). Check for floor/ceiling effects.',
   'Trung bình, SD, độ lệch, độ nhọn, và liên tương quan của '
   '10 chỉ báo (Bảng 1). Kiểm tra hiệu ứng sàn/trần.')

H2('3.2 — LPA model selection')
PB('Table 2 reporting BIC, AIC, entropy, LMR-LRT p-value, '
   'and smallest class % for K = 2, 3, 4, 5, 6 solutions. '
   'Convergence checks. Final K determined by triangulation '
   'of criteria + interpretability — exact K to be determined '
   'by analysis.',
   'Bảng 2 báo cáo BIC, AIC, entropy, giá trị p LMR-LRT, '
   'và % lớp nhỏ nhất cho các giải pháp K = 2, 3, 4, 5, 6. '
   'Kiểm tra hội tụ. K cuối cùng xác định bằng triangulation '
   'các tiêu chí + tính diễn giải — K chính xác chờ phân tích '
   'xác định.')

H2('3.3 — Profile descriptions')
PB('Figure 1 (radar chart or line graph) showing standardized '
   'mean profile across 10 indicators for each profile. Table '
   '3 with descriptive label + prevalence + indicator-specific '
   'means + SDs.',
   'Hình 1 (biểu đồ radar hoặc đồ thị đường) thể hiện trung '
   'bình chuẩn hóa của hồ sơ qua 10 chỉ báo cho mỗi hồ sơ. '
   'Bảng 3 với nhãn mô tả + tỷ lệ + trung bình + SD theo chỉ '
   'báo.')

H2('3.4 — External validation: gender, grade, school')
PB('Chi-square tests of profile membership × gender, × grade, '
   '× school. Report cell %, residuals, Cohen\'s w. Discuss '
   'any profile that disproportionately contains one '
   'sub-population.',
   'Kiểm định chi-square thành viên hồ sơ × giới, × khối lớp, '
   '× trường. Báo cáo % ô, phần dư, Cohen\'s w. Bàn luận bất '
   'kỳ hồ sơ nào chứa lệch một dân số con.')

H2('3.5 — Profile-anxiety severity comparison')
PB('ANOVA comparing GAD, SAD, SocAD raw scores across '
   'profiles. Post-hoc Tukey comparisons. Eta-squared effect '
   'sizes.',
   'ANOVA so sánh điểm thô GAD, SAD, SocAD giữa các hồ sơ. '
   'So sánh post-hoc Tukey. Effect size eta-squared.')


# =====================================================
H1('4. DISCUSSION (target 1,500-1,800 words full draft)')

H2('4.1 — Summary of profiles found')
PB('Brief recap: K profiles identified, named descriptively, '
   'with their prevalences and key characteristics.',
   'Tóm tắt ngắn: K hồ sơ được nhận diện, đặt tên mô tả, kèm '
   'tỷ lệ và đặc điểm chính.')

H2('4.2 — Theoretical contribution: complement to SEM (Q2/Q3)')
PB('SEM tells us WHAT pathways matter (Q2) and FOR WHOM '
   'differently (Q3). LPA tells us HOW students are clustered '
   'into typologies. The three perspectives together provide '
   'a comprehensive picture of Vietnamese adolescent anxiety '
   'that no single method achieves. Discuss how variable-'
   'centered + person-centered findings can converge or '
   'diverge (Marsh et al. 2009).',
   'SEM cho biết các đường dẫn nào quan trọng (Q2) và CHO AI '
   'khác nhau (Q3). LPA cho biết học sinh được gộp thành các '
   'phân loại như thế nào. Ba góc nhìn cùng nhau cung cấp '
   'bức tranh toàn diện về lo âu vị thành niên Việt Nam mà '
   'không phương pháp đơn lẻ nào đạt được. Bàn luận cách các '
   'phát hiện lấy biến + lấy con người làm trung tâm có thể '
   'hội tụ hoặc phân kỳ (Marsh và cs. 2009).')

H2('4.3 — Co-rumination and gender-role implications')
PB('Some profiles may show distinctive gender distributions '
   'consistent with co-rumination patterns (Rose 2002) and '
   'gender-role intensification at this developmental window '
   '(Hill & Lynch 1983; Hankin et al. 2007; McLean & Anderson '
   '2009). For example, a "peer-victimization + low-school-'
   'membership" profile may over-represent girls, paralleling '
   'our Q3 findings on bullying → SAD pathway non-invariance.',
   'Một số hồ sơ có thể thể hiện phân bố giới đặc trưng nhất '
   'quán với mẫu hình co-rumination (Rose 2002) và sự gia '
   'tăng vai trò giới ở cửa sổ phát triển này (Hill & Lynch '
   '1983; Hankin và cs. 2007; McLean & Anderson 2009). Ví dụ '
   'hồ sơ "nạn nhân hóa từ bạn bè + thấp gắn bó trường" có '
   'thể đại diện quá mức nữ, song song với phát hiện Q3 về '
   'không bất biến đường dẫn bắt nạt → SAD.')

H2('4.4 — Cultural-developmental interpretation')
PB('Reuse cultural framing from Q2 v1 Section 4.4 (tam giao + '
   'Confucian + co-rumination) — interpret which profiles '
   'reflect typical Vietnamese socio-cultural niches. For '
   'example, "high-academic-stress" profile likely reflects '
   'the Confucian achievement script (Stankov 2010); "high-'
   'peer-victimization + low-school-membership" reflects '
   'breakdown of the Confucian harmonious school ideal.',
   'Sử dụng lại khung văn hóa từ Q2 v1 Section 4.4 (tam giáo + '
   'Nho giáo + co-rumination) — diễn giải hồ sơ nào phản ánh '
   'các vị thế xã hội-văn hóa Việt Nam điển hình. Ví dụ, hồ sơ '
   '"cao áp lực học tập" có thể phản ánh kịch bản thành tích '
   'Nho giáo (Stankov 2010); "cao nạn nhân hóa từ bạn bè + '
   'thấp gắn bó trường" phản ánh sự đổ vỡ của lý tưởng trường '
   'học hài hòa Nho giáo.')

H2('4.5 — Practical implications: profile-tailored intervention')
PB('Each profile suggests a different intervention package. '
   'Resilient majority: universal psycho-education on stress '
   'management. Achievement-driven: cognitive-behavioral '
   'reframing of academic pressure + self-compassion. Peer-'
   'victimization: anti-bullying program + assertiveness '
   'training. Digital-isolation: parental engagement program + '
   'screen-time management. Severe-anxiety: clinical referral.',
   'Mỗi hồ sơ gợi ý một gói can thiệp khác nhau. Đa số kiên '
   'cường: giáo dục tâm lý phổ quát về quản lý stress. Theo '
   'đuổi thành tích: tái cấu trúc nhận thức-hành vi về áp lực '
   'học tập + tự thương xót. Nạn nhân hóa từ bạn bè: chương '
   'trình chống bắt nạt + huấn luyện quyết đoán. Cô lập số: '
   'chương trình kết nối cha mẹ + quản lý thời gian màn hình. '
   'Lo âu nặng: chuyển tuyến lâm sàng.')

H2('4.6 — Limitations')
PB('(a) Cross-sectional design: profiles are snapshots, '
   'transitions between profiles over time cannot be tested; '
   '(b) LPA assumes profile homogeneity — within-profile '
   'variability may still be substantial; (c) profile '
   'replicability requires external samples; (d) LMR-LRT and '
   'BIC may disagree — model selection has irreducible '
   'subjectivity.',
   '(a) Thiết kế cắt ngang: hồ sơ là ảnh chụp nhanh, chuyển '
   'tiếp giữa các hồ sơ theo thời gian không thể kiểm chứng; '
   '(b) LPA giả định tính đồng nhất trong hồ sơ — biến thiên '
   'trong-hồ-sơ có thể vẫn đáng kể; (c) khả năng tái tạo hồ '
   'sơ yêu cầu các mẫu bên ngoài; (d) LMR-LRT và BIC có thể '
   'không đồng ý — chọn mô hình có tính chủ quan không thể '
   'loại bỏ.')

H2('4.7 — Future directions')
PB('(1) Latent transition analysis (LTA) on a longitudinal '
   'extension to test profile stability over time; (2) '
   'replicate in rural + minority-ethnic cohorts; (3) link '
   'profiles to objective outcomes (school attendance, '
   'grades, healthcare utilization).',
   '(1) Phân tích chuyển tiếp tiềm ẩn (LTA) trên phần mở rộng '
   'dọc để kiểm chứng độ ổn định hồ sơ theo thời gian; (2) '
   'lặp lại ở cohort nông thôn + dân tộc thiểu số; (3) liên '
   'kết hồ sơ với các kết quả khách quan (chuyên cần, học '
   'lực, sử dụng dịch vụ y tế).')


# =====================================================
H1('5. REFERENCES (target 30-40 entries full draft)')

NOTE('12 KEY references shared with Q2/Q3 already verified via PubMed. '
     'Q4-specific methodology refs (Bauer 2007 LCA; Lanza & Cooper '
     '2016 LPA; Marsh et al. 2009 person- vs variable-centered; '
     'Nylund et al. 2007 LMR-LRT) to be added in full draft.',
     '12 tham khảo CHỦ CHỐT chia sẻ với Q2/Q3 đã verify qua PubMed. '
     'Các tham khảo phương pháp luận đặc thù Q4 (Bauer 2007 LCA; '
     'Lanza & Cooper 2016 LPA; Marsh và cs. 2009 lấy con người vs '
     'lấy biến; Nylund và cs. 2007 LMR-LRT) sẽ thêm vào draft đầy đủ.')

refs = [
    '[1] Hankin BL, Mermelstein R, Roesch L. (2007). Sex '
    'differences in adolescent depression. Child Development '
    '78(1):279-295. PMID: 17328705.',
    '[2] McLean CP, Anderson ER. (2009). Brave men and timid '
    'women? Clinical Psychology Review 29(6):496-505. '
    'PMID: 19541399.',
    '[3] Kessler RC, Berglund P, Demler O, et al. (2005). '
    'Lifetime prevalence and age-of-onset distributions of '
    'DSM-IV disorders in the NCS-R. Archives of General '
    'Psychiatry 62(6):593-602. PMID: 15939837.',
    '[4] Rose AJ. (2002). Co-rumination in the friendships of '
    'girls and boys. Child Development 73(6):1830-1843. '
    'PMID: 12487497.',
    '[5] Stankov L. (2010). Unforgiving Confucian culture. '
    'Learning and Individual Differences 20(6):555-563.',
    '[6] Small S, Blanc J. (2021). Mental Health During COVID-19: '
    'Tam Giao and Vietnam\'s Response. Frontiers in Psychiatry '
    '11:589618. PMID: 33536961.',
    '[7] Bauer DJ. (2007). Observations on the use of growth '
    'mixture models in psychological research. Multivariate '
    'Behavioral Research 42(4):757-786. [TO VERIFY PMID in '
    'full draft.]',
    '[8] Lanza ST, Cooper BR. (2016). Latent class analysis for '
    'developmental research. Child Development Perspectives '
    '10(1):59-64.',
    '[9] Marsh HW, Lüdtke O, Trautwein U, Morin AJS. (2009). '
    'Classical latent profile analysis of academic self-concept '
    'dimensions. Structural Equation Modeling 16(2):191-225.',
    '[10] Nylund KL, Asparouhov T, Muthén BO. (2007). Deciding '
    'on the number of classes in latent class analysis and '
    'growth mixture modeling. Structural Equation Modeling '
    '14(4):535-569.',
    '[11] Compas BE, Jaser SS, Bettis AH, et al. (2017). Coping, '
    'emotion regulation, and psychopathology in childhood and '
    'adolescence. Psychological Bulletin 143(9):939-991. '
    'PMID: 28616996.',
    '[12] Jefferies P, Ungar M. (2020). Social anxiety in young '
    'people. PLOS ONE 15(9):e0239133. DOI: '
    '10.1371/journal.pone.0239133.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Footer
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Soạn 08/06/2026 — Outline Q4 v1, kế hoạch năm 2027, chờ '
              'cả nhóm cùng quyết định')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
