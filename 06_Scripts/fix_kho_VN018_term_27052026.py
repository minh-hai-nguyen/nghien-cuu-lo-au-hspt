# -*- coding: utf-8 -*-
"""Fix kho - VN018 author order + global term 'tổng quát' -> 'lan tỏa' in all ban-dich + tom-tat.
27/05/2026."""
import os, sys, io, glob
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
RED = RGBColor(192, 0, 0)

# VN018 author order fixes
VN018_FILES = [
    '03_Ban-dich/VN018_AnGiang_2025_YHVN.docx',
    'Tom-tat-tung-bai/VN018_AnGiang_2025_YHVN.docx',
]
VN018_REPLACEMENTS = [
    ('Lê Minh T., Nguyễn Đăng K., Ngô Anh V.', 'Nguyễn Đăng K., Lê Minh T., Ngô Anh V.'),
    ('Lê Minh T., Nguyễn Đăng K. và Ngô Anh V.', 'Nguyễn Đăng K., Lê Minh T. và Ngô Anh V.'),
    ('Lê Minh T., Nguyễn Đăng Khoa, Ngô Anh Vinh', 'Nguyễn Đăng Khoa, Lê Minh Thi, Ngô Anh Vinh'),
    ('Lê Minh Thi, Nguyễn Đăng Khoa, & Ngô Anh Vinh', 'Nguyễn Đăng Khoa, Lê Minh Thi, & Ngô Anh Vinh'),
    ('Lê Minh Thi, Nguyễn Đăng Khoa và Ngô Anh Vinh', 'Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh'),
]

# Global term replacements - applied to ALL ban-dich + tom-tat
TERM_REPLACEMENTS = [
    ('Rối loạn Lo âu Tổng quát', 'Rối loạn Lo âu Lan tỏa'),
    ('Rối loạn lo âu tổng quát', 'Rối loạn lo âu lan tỏa'),
    ('rối loạn lo âu tổng quát', 'rối loạn lo âu lan tỏa'),
    ('Lo âu Tổng quát', 'Lo âu Lan tỏa'),
    ('Lo âu tổng quát', 'Lo âu lan tỏa'),
    ('lo âu tổng quát', 'lo âu lan tỏa'),
]


def apply_replacements_in_paragraph(p, replacements, applied_tracker, label):
    full = p.text
    if not full:
        return False
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            idx = full.find(old, cursor)
            if idx == -1: break
            matches.append((idx, idx + len(old), old, new))
            cursor = idx + len(old)
    if not matches:
        return False
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    safe_matches = []
    last_end = -1
    for m in matches:
        if m[0] >= last_end:
            safe_matches.append(m)
            last_end = m[1]
    if not safe_matches:
        return False
    segments = []
    cursor = 0
    for start, end, old, new in safe_matches:
        if cursor < start:
            segments.append((full[cursor:start], False))
        segments.append((new, True))
        applied_tracker.append(f"{label}: '{old[:50]}' -> '{new[:50]}'")
        cursor = end
    if cursor < len(full):
        segments.append((full[cursor:], False))
    for r in p.runs:
        r.text = ''
    for text, is_red in segments:
        if not text:
            continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


def fix_doc(relpath, replacements):
    full = os.path.join(ROOT, relpath)
    if not os.path.exists(full):
        return None, 0, []
    d = Document(full)
    applied = []
    # Paragraphs
    for i, p in enumerate(d.paragraphs):
        apply_replacements_in_paragraph(p, replacements, applied, f"para{i}")
    # Tables
    for ti, tb in enumerate(d.tables):
        for ri, row in enumerate(tb.rows):
            for ci, cell in enumerate(row.cells):
                for cpi, p in enumerate(cell.paragraphs):
                    apply_replacements_in_paragraph(p, replacements, applied,
                                                     f"T{ti+1}r{ri}c{ci}p{cpi}")
    if applied:
        base, ext = os.path.splitext(full)
        out = base + '_FIXED_27052026' + ext
        d.save(out)
        return out, len(applied), applied
    return None, 0, []


print("=" * 70)
print("PHASE 1: Fix VN018 author order")
print("=" * 70)
for relpath in VN018_FILES:
    # Combine VN018 replacements + term replacements
    all_rep = VN018_REPLACEMENTS + TERM_REPLACEMENTS
    out, n, applied = fix_doc(relpath, all_rep)
    print(f"\n--- {relpath} ---")
    if out:
        print(f"  Output: {out}")
        print(f"  Applied {n} fixes:")
        for a in applied:
            print(f"    ✓ {a}")
    else:
        print(f"  No fixes applied (or file not found)")

print("\n" + "=" * 70)
print("PHASE 2: Fix term 'tổng quát' -> 'lan tỏa' in all ban-dich + tom-tat")
print("=" * 70)

ban_dich_files = glob.glob(os.path.join(ROOT, '03_Ban-dich', '*.docx'))
tom_tat_files = glob.glob(os.path.join(ROOT, 'Tom-tat-tung-bai', '*.docx'))
all_files = ban_dich_files + tom_tat_files

print(f"Total files: {len(all_files)}")

term_summary = {}
for full in all_files:
    relpath = os.path.relpath(full, ROOT).replace('\\', '/')
    # Skip VN018 (already handled with combined fix)
    if 'VN018_AnGiang' in relpath or 'VN017_NguyenDanhLam_2022_YHVN_FIXED' in relpath:
        continue
    if '_FIXED_' in relpath:  # Skip already-fixed files
        continue
    out, n, applied = fix_doc(relpath, TERM_REPLACEMENTS)
    if out:
        term_summary[relpath] = (out, n, applied)

print(f"\nTotal files with term fixes: {len(term_summary)}")
for rel, (out, n, applied) in term_summary.items():
    print(f"\n  {rel}: {n} fixes")
    for a in applied[:3]:
        print(f"    ✓ {a}")
    if len(applied) > 3:
        print(f"    ... +{len(applied)-3} more")
