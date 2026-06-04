# -*- coding: utf-8 -*-
"""Dịch đầy đủ QT28 — AJP 2024 — Zugman et al. — Pediatric Anxiety Treatment"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link bài báo gốc: https://doi.org/10.1176/appi.ajp.20231037', size=10)

add_heading(doc, 'Phương pháp tiếp cận hiện tại và tương lai trong điều trị rối loạn lo âu ở trẻ em', 1)
h = doc.add_paragraph()
r = h.add_run('Current and Future Approaches to Pediatric Anxiety Disorder Treatment')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Current and Future Approaches to Pediatric Anxiety Disorder Treatment'),
    ('Tiêu đề dịch', 'Phương pháp tiếp cận hiện tại và tương lai trong điều trị rối loạn lo âu ở trẻ em'),
    ('Tác giả', 'Andre Zugman, MD, PhD; Anderson M. Winkler, MD, PhD; Purnima Qamar, BS; Daniel S. Pine, MD'),
    ('Cơ quan', 'Section on Development and Affective Neuroscience, National Institute of Mental Health (NIMH), National Institutes of Health (NIH), Bethesda, MD, Hoa Kỳ'),
    ('Tạp chí', 'American Journal of Psychiatry — AJP (Q1, IF ≈ 18)'),
    ('Thông tin xuất bản', '2024, Vol. 181, No. 3, pp. 189–200, 12 trang'),
    ('DOI', '10.1176/appi.ajp.20231037'),
    ('Loại tài liệu', 'Tổng quan (Overview) — đánh giá phê bình tài liệu'),
    ('Phạm vi', '3 rối loạn lo âu trẻ em: Lo âu Tổng quát (GAD), Lo âu Chia ly, Lo âu Xã hội'),
])
add_page_ref(doc, '189–200', 'American Journal of Psychiatry', 'Vol. 181(3), March 2024')

# TÓM TẮT
add_heading(doc, 'TÓM TẮT', 2)
p = doc.add_paragraph()
r = p.add_run('Tổng quan này đánh giá phê bình tài liệu về điều trị rối loạn lo âu trẻ em. Hai phương pháp điều trị đã thiết lập: liệu pháp nhận thức–hành vi (CBT) và thuốc chống trầm cảm (SSRIs). Nhiều trẻ nhận điều trị này không đạt thuyên giảm — tạo nhu cầu phương pháp mới. Sau khi tóm tắt CBT và thuốc hiện có, tác giả mô tả NC đặt nền tảng cho cải thiện điều trị, tận dụng các khám phá thần kinh khoa học cung cấp cái nhìn sâu về cơ chế điều trị thành công.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

add_p(doc, 'Bài viết tập trung vào 3 rối loạn lo âu trẻ em: rối loạn lo âu tổng quát (GAD), rối loạn lo âu chia ly (separation anxiety), và rối loạn lo âu xã hội (social anxiety). Xem xét cùng nhóm vì hầu hết NC điều trị nhắm vào chúng cùng nhau (tỷ lệ cao, đi kèm thường xuyên). Không bao gồm phobia cụ thể (tiên lượng tốt hơn) và rối loạn hoảng sợ (hiếm trước tuổi trưởng thành).')

# ĐÁNH GIÁ NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'AJP Q1 IF ≈ 18 — tạp chí tâm thần hàng đầu thế giới. Tác giả từ NIMH/NIH — cơ quan NC SKTT hàng đầu Mỹ.',
    'Tổng quan toàn diện: CBT + SSRIs + kết hợp + hướng tương lai (ABMT, DCS, VR, biomarker).',
    'CAMS (Child/Adolescent Anxiety Multimodal Study): RCT LỚN NHẤT — kết hợp CBT+sertraline đạt 80,7% đáp ứng vs CBT đơn 59,7% vs SSRI đơn 54,9% vs placebo 23,7%.',
    'NNT (Number Needed to Treat) = 4 cho SSRIs — hiệu quả mạnh.',
    'CBT: NNT ≈ 3; phơi nhiễm (exposure) là thành phần cốt lõi; hiệu quả duy trì 6–12 tháng.',
    'Thần kinh khoa học: amygdala, mạch đe dọa, attention bias — nền tảng cho phương pháp mới.',
    'Vấn đề: 20–40% KHÔNG đáp ứng điều trị → cần đổi mới.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Tổng quan narrative — KHÔNG phải phân tích tổng hợp hệ thống (systematic review/meta-analysis).',
    'Chủ yếu dữ liệu phương Tây — ít đề cập châu Á/LMIC. BMC NMA 2025 (QT29) bao gồm thêm PE, VRET, ACT.',
    'Hướng tương lai (ABMT, DCS, VR) chưa có ĐỦ bằng chứng RCT.',
    'Tập trung 3 loại lo âu — không bao gồm phobia cụ thể, PTSD.',
]:
    add_p(doc, f'• {b}')

# 1. ĐIỀU TRỊ ĐÃ THIẾT LẬP
add_page_ref(doc, '189–194', 'AJP', 'Vol. 181(3), 2024')
add_heading(doc, '1. ĐIỀU TRỊ ĐÃ THIẾT LẬP', 2)

add_p(doc, '1.1. Dược lý — SSRIs', bold=True)
add_p(doc, 'Thuốc ức chế tái hấp thu serotonin chọn lọc (SSRIs) là điều trị dược lý HÀNG ĐẦU. Phân tích tổng hợp cung cấp bằng chứng mạnh mẽ về hiệu quả cho hầu hết tất cả SSRIs, với NNT (Number Needed to Treat) ≈ 4 cho đáp ứng (response). Nghĩa là: cứ 4 trẻ điều trị bằng SSRI, 1 trẻ đáp ứng nhờ thuốc (so với placebo).')
add_p(doc, 'Các SSRI hiệu quả: fluoxetine, sertraline, fluvoxamine — hiệu quả tương đương. FDA phê duyệt: fluoxetine (trầm cảm trẻ 8+), fluvoxamine (OCD trẻ 8+), sertraline (OCD trẻ 6+). Mặc dù không phê duyệt riêng cho lo âu, SSRIs được sử dụng rộng rãi off-label.')
add_p(doc, 'SNRIs (Serotonin-Norepinephrine Reuptake Inhibitors): duloxetine, venlafaxine — cũng hiệu quả nhưng ít dữ liệu hơn, nhiều tác dụng phụ hơn.')
add_p(doc, 'FDA cảnh báo hộp đen: TẤT CẢ SSRIs/SNRIs có cảnh báo về nguy cơ tự tử ở VTN. Tuy nhiên, phân tích tổng hợp cho thấy nguy cơ THẤP và LỢI ÍCH vượt trội rủi ro. Tỷ lệ tự tử: NNH (Number Needed to Harm) ≈ 143 (nghĩa là cứ 143 trẻ điều trị, 1 trẻ có thể có ý tưởng tự tử do thuốc). So với NNT = 4 → lợi ích >> rủi ro.')

add_p(doc, 'Thuốc thay thế (hàng 2–3):', bold=True)
add_p(doc, '• Tricyclics (TCAs): KÉM HIỆU QUẢ hơn SSRIs, nhiều tác dụng phụ hơn, nguy hiểm hơn khi quá liều → hàng 3.')
add_p(doc, '• Benzodiazepines: chỉ 3 RCT NHỎ ở trẻ em (lớn nhất: 30 bệnh nhân alprazolam). Không có bằng chứng hiệu quả trong phân tích tổng hợp. Nguy cơ phụ thuộc. → Không khuyến nghị.')
add_p(doc, '• Buspirone: ít dữ liệu. 1 RCT trẻ 6–17 (n = 559): không hiệu quả hơn placebo. → Không khuyến nghị.')
add_p(doc, '• Gabapentin: ít bằng chứng ở trẻ em.')

add_p(doc, '1.2. CBT — Liệu pháp Nhận thức–Hành vi', bold=True)
add_p(doc, 'CBT là điều trị tâm lý HÀNG ĐẦU. Phân tích tổng hợp cho thấy hiệu quả trung bình–lớn (NNT ≈ 3 cho đáp ứng). Thành phần cốt lõi:')
add_p(doc, '• Giáo dục tâm lý (psychoeducation): giải thích lo âu, mô hình nhận thức–hành vi.')
add_p(doc, '• Tái cấu trúc nhận thức (cognitive restructuring): nhận diện và thay đổi suy nghĩ tiêu cực tự động.')
add_p(doc, '• PHƠI NHIỄM DẦN (graduated exposure): thành phần QUAN TRỌNG NHẤT — tiếp xúc dần với tình huống sợ hãi trong môi trường an toàn. Tạo "học tập ức chế" (inhibitory learning) — hình thành liên kết mới cạnh tranh với ký ức sợ hãi.')
add_p(doc, '• Thư giãn, quản lý lo lắng, củng cố tích cực.')
add_p(doc, 'Hiệu quả duy trì 6–12 tháng sau điều trị (NC theo dõi). Tuy nhiên, 30–40% KHÔNG đáp ứng đầy đủ → cần cải thiện.')

add_p(doc, '1.3. Kết hợp CBT + SSRI', bold=True)
add_p(doc, 'CAMS (Child/Adolescent Anxiety Multimodal Study) — RCT LỚN NHẤT và quan trọng nhất:')

add_heading(doc, 'Bảng 1. Kết quả CAMS — Kết hợp CBT + SSRI cho lo âu trẻ em', 3)
add_table(doc,
    ['Phương pháp', 'Tỷ lệ đáp ứng', 'NNT vs placebo', 'Ghi chú'],
    [['CBT + Sertraline (kết hợp)', '80,7%', '~2', 'HIỆU QUẢ NHẤT'],
     ['CBT đơn', '59,7%', '~3', 'Phơi nhiễm là thành phần chính'],
     ['Sertraline đơn', '54,9%', '~4', 'Cảnh báo FDA tự tử'],
     ['Placebo', '23,7%', '—', 'Hiệu ứng placebo đáng kể']],
    widths=[4.0, 2.5, 2.5, 4.5])
add_p(doc, 'CAMS: N = 488 trẻ 7–17 tuổi, 3 loại lo âu (GAD, SAD, social anxiety). 12 tuần điều trị. Đáp ứng = CGI-I ≤ 2. Kết hợp VƯỢT TRỘI cả đơn trị (Walkup et al. 2008, NEJM).', size=9, italic=True)

add_heading(doc, 'Bảng 2. Tổng hợp phân tích tổng hợp gần đây về điều trị lo âu trẻ em', 3)
add_table(doc,
    ['Phân tích tổng hợp', 'Phương pháp', 'Số NC/Mẫu', 'Phát hiện chính'],
    [['Strawn et al. 2023', 'SSRIs/SNRIs', '22 RCT, ~3.500', 'NNT=4; tất cả SSRIs hiệu quả; SNRIs cũng nhưng tác dụng phụ nhiều hơn'],
     ['James et al. 2020 (Cochrane)', 'CBT vs đối chứng', '87 NC, ~5.900', 'CBT hiệu quả; NNT≈3; phơi nhiễm quan trọng nhất'],
     ['Walkup et al. 2008 (CAMS)', 'CBT+SSRI vs đơn', '488 trẻ, 1 RCT', 'Kết hợp 80,7% > CBT 59,7% > SSRI 54,9% > placebo 23,7%'],
     ['Dobson et al. 2019', 'Duy trì dài hạn', 'Nhiều NC', 'CBT hiệu quả ≥6 tháng; SSRI cần duy trì lâu hơn'],
     ['Li et al. 2025 (QT29, NMA)', 'ACT/CBT/PE/VRET', '30 RCT, 1.711', 'ACT hạng 1 (SUCRA 0,69); CBT hạng 2 (0,66); PE+VRET cũng hiệu quả']],
    widths=[3.0, 2.5, 2.5, 5.5])

# 2. HƯỚNG CẢI THIỆN NGẮN HẠN
add_page_ref(doc, '194–197', 'AJP', 'Vol. 181(3), 2024')
add_heading(doc, '2. HƯỚNG CẢI THIỆN NGẮN HẠN', 2)

add_p(doc, '2.1. Attention Bias Modification Training (ABMT)', bold=True)
add_p(doc, 'Huấn luyện chú ý: tập luyện hệ thống để "chuyển" chú ý TRÁNH xa kích thích đe dọa. Dựa trên phát hiện: người lo âu có THIÊN LỆCH CHÚ Ý hướng đến đe dọa (attention bias to threat) — liên quan hoạt động amygdala tăng cường. ABMT sử dụng bài tập dot-probe để huấn luyện chú ý tránh mặt sợ hãi/giận dữ.')
add_p(doc, 'Kết quả LẪN LỘN: một số NC cho thấy hiệu quả nhỏ–trung bình, nhưng nhiều NC không tìm thấy. Phân tích tổng hợp gần đây (Cristea et al. 2015): hiệu ứng NHỎ. Tuy nhiên, kết hợp ABMT + CBT có thể hiệu quả hơn đơn lẻ — "tăng cường" CBT bằng chuẩn bị chú ý trước phơi nhiễm.')

add_p(doc, '2.2. D-cycloserine (DCS)', bold=True)
add_p(doc, 'DCS là kháng sinh có tác dụng phụ: tăng cường glutamate tại receptor NMDA → tăng cường consolidation (củng cố) ký ức mới trong học tập phơi nhiễm. Lý thuyết: nếu CBT hoạt động qua "học tập ức chế", DCS có thể tăng cường quá trình này.')
add_p(doc, 'Kết quả ở người lớn: hiệu quả NHỎ nhưng đáng kể (phân tích tổng hợp Mataix-Cols et al. 2017). Ở trẻ em: RẤT ÍT dữ liệu — 2 RCT nhỏ, kết quả lẫn lộn. Cần thêm NC.')

add_p(doc, '2.3. Thực tế ảo (Virtual Reality — VR)', bold=True)
add_p(doc, 'VR cho phép phơi nhiễm trong môi trường ảo — thuận tiện (không cần đến tình huống thực), kiểm soát được (điều chỉnh mức độ kích thích), lặp lại được. Đặc biệt hữu ích cho phobia cụ thể và lo âu xã hội.')
add_p(doc, 'Tuy nhiên, ÍT RCT ở trẻ em. BMC NMA 2025 (QT29, Li et al.): VRET xếp hạng 3 (SUCRA 0,51) — hiệu quả nhưng sau ACT và CBT. Cần thêm NC với mẫu lớn hơn.')

add_p(doc, '2.4. Can thiệp dựa trên internet/số (digital)', bold=True)
add_p(doc, 'CBT trực tuyến (iCBT) — tiềm năng mở rộng tiếp cận, đặc biệt ở vùng thiếu nhân lực. Tuy nhiên, hiệu quả thường thấp hơn CBT trực tiếp, tỷ lệ bỏ cuộc cao hơn. Cần nghiên cứu thêm về tối ưu hóa.')

# 3. HƯỚNG ĐỔI MỚI DÀI HẠN
add_page_ref(doc, '197–199', 'AJP', 'Vol. 181(3), 2024')
add_heading(doc, '3. HƯỚNG ĐỔI MỚI DÀI HẠN — THẦN KINH KHOA HỌC', 2)

add_p(doc, '3.1. Mạch thần kinh đe dọa', bold=True)
add_p(doc, 'Lo âu liên quan đến rối loạn chức năng trong mạch thần kinh phản ứng đe dọa ở động vật có vú — BẢO TỒN XUYÊN LOÀI (từ chuột đến người). Amygdala — trung tâm xử lý đe dọa — hoạt động TĂNG CƯỜNG ở trẻ lo âu khi tiếp xúc khuôn mặt giận dữ/sợ hãi (fMRI). Vỏ não trước trán (prefrontal cortex) — điều tiết amygdala — chức năng yếu hơn ở trẻ lo âu.')
add_p(doc, 'Hàm ý: can thiệp nhắm vào TĂNG CƯỜNG điều tiết trước trán → amygdala có thể hiệu quả. CBT hoạt động một phần qua cơ chế này (phơi nhiễm → "học tập ức chế" → tăng kiểm soát trước trán lên amygdala).')

add_p(doc, '3.2. Biomarker dự báo đáp ứng điều trị', bold=True)
add_p(doc, 'Nếu xác định TRƯỚC ai đáp ứng CBT vs SSRI: cá nhân hóa điều trị (precision medicine). Ứng viên biomarker:')
add_p(doc, '• fMRI amygdala phản ứng đe dọa: hoạt động amygdala CAO trước điều trị → đáp ứng CBT TỐT hơn (có thể vì có nhiều "room for improvement").')
add_p(doc, '• EEG: Error-Related Negativity (ERN) — trẻ lo âu có ERN tăng (phát hiện lỗi quá mức). ERN có thể dự báo đáp ứng.')
add_p(doc, '• Cortisol, CRP, interleukins — biomarker máu nhạy cảm với đe dọa.')

add_heading(doc, 'Bảng 3. Tóm tắt hướng cải thiện điều trị lo âu trẻ em', 3)
add_table(doc,
    ['Hướng', 'Phương pháp', 'Bằng chứng', 'Ứng dụng VN'],
    [['ABMT', 'Huấn luyện chú ý tránh đe dọa', 'Lẫn lộn; NHỎ đơn lẻ; hứa hẹn kết hợp CBT', 'Có thể số hóa (app)'],
     ['DCS + CBT', 'Thuốc tăng cường học tập phơi nhiễm', 'Hiệu quả ở người lớn; ÍT dữ liệu trẻ em', 'Cần RCT ở VN'],
     ['VR/VRET', 'Phơi nhiễm qua thực tế ảo', 'SUCRA 0,51 (NMA QT29); ít RCT trẻ em', 'Cần công nghệ + chi phí'],
     ['iCBT (trực tuyến)', 'CBT qua internet', 'Hiệu quả thấp hơn trực tiếp; bỏ cuộc cao', 'Phù hợp VN (thiếu nhân lực)'],
     ['Biomarker (fMRI)', 'Dự báo đáp ứng cá nhân', 'NC sớm; chưa áp dụng lâm sàng', 'Tương lai xa'],
     ['ACT (QT29)', 'Chấp nhận + cam kết giá trị', 'SUCRA 0,69 — hạng 1 NMA', 'Mới ở VN; cần đào tạo'],
     ['PE (QT29)', 'Hoạt động thể chất', 'SUCRA 0,51; endorphin, BDNF', 'DỄ triển khai tại trường VN']],
    widths=[2.5, 3.5, 4.0, 3.5])

# 4. KẾT LUẬN
add_heading(doc, '4. KẾT LUẬN', 2)
add_p(doc, 'CBT và SSRIs là điều trị đã thiết lập cho rối loạn lo âu trẻ em. Kết hợp CBT+SSRI hiệu quả nhất (CAMS: 80,7% đáp ứng). Phơi nhiễm là thành phần CBT quan trọng nhất. SSRIs: NNT = 4, an toàn mặc dù cảnh báo FDA. Tuy nhiên, 20–40% KHÔNG đáp ứng → cần cải thiện: ABMT, DCS, VR, iCBT, biomarker.')
add_p(doc, 'NC thần kinh khoa học (amygdala, attention bias, ERN) đặt nền tảng cho đổi mới nhưng cần nhiều RCT hơn. Đặc biệt thiếu dữ liệu ở LMIC/châu Á (bao gồm VN). BMC NMA 2025 (QT29): ACT và PE cũng hiệu quả — mở rộng lựa chọn ngoài CBT+SSRIs.')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
refs = [
    'Zugman, A., Winkler, A.M., Qamar, P. & Pine, D.S. (2024). Current and future approaches to pediatric anxiety disorder treatment. Am J Psychiatry, 181, 189–200.',
    'Walkup, J.T., et al. (2008). Cognitive behavioral therapy, sertraline, or a combination in childhood anxiety. NEJM, 359, 2753–2766.',
    'Strawn, J.R., et al. (2023). Pharmacotherapy for pediatric anxiety disorders. Pediatrics, 152(3), e2023063285.',
    'James, A.C., et al. (2020). Cognitive behavioural therapy for anxiety disorders in children and adolescents. Cochrane Database Syst Rev, 11, CD004690.',
    'Cristea, I.A., et al. (2015). A critical examination of attentional bias modification for anxiety disorders. BMJ, 351, h4919.',
    'Mataix-Cols, D., et al. (2017). D-cycloserine augmentation of exposure-based cognitive behavior therapy for anxiety. JAMA Psychiatry, 74, 501–510.',
    'Li, L.H., et al. (2025). Effectiveness of different types of interventions for anxiety disorders in children and adolescents: NMA. BMC Psychiatry, 25, 809.',
]
for ref in refs:
    add_p(doc, ref, size=10)
add_p(doc, '(Xem đầy đủ trong bài gốc — 90+ TLTK)', size=10, italic=True)

# VIẾT TẮT
add_abbreviation_table(doc, [
    ('CBT', 'Cognitive-Behavioral Therapy — Liệu pháp Nhận thức–Hành vi'),
    ('SSRI', 'Selective Serotonin Reuptake Inhibitor — Thuốc Ức chế Tái hấp thu Serotonin Chọn lọc'),
    ('SNRI', 'Serotonin-Norepinephrine Reuptake Inhibitor'),
    ('CAMS', 'Child/Adolescent Anxiety Multimodal Study'),
    ('NNT', 'Number Needed to Treat — Số Bệnh nhân Cần Điều trị để 1 người hưởng lợi'),
    ('NNH', 'Number Needed to Harm — Số Bệnh nhân Điều trị để 1 người bị hại'),
    ('ABMT', 'Attention Bias Modification Training — Huấn luyện Điều chỉnh Thiên lệch Chú ý'),
    ('DCS', 'D-cycloserine — Thuốc tăng cường học tập phơi nhiễm'),
    ('VR/VRET', 'Virtual Reality Exposure Therapy — Liệu pháp Phơi nhiễm Thực tế Ảo'),
    ('iCBT', 'Internet-delivered CBT — CBT qua Internet'),
    ('ACT', 'Acceptance and Commitment Therapy — Liệu pháp Chấp nhận và Cam kết'),
    ('PE', 'Physical Exercise — Hoạt động Thể chất'),
    ('NIMH', 'National Institute of Mental Health — Viện SKTT Quốc gia Mỹ'),
    ('NIH', 'National Institutes of Health — Viện Y tế Quốc gia Mỹ'),
    ('FDA', 'Food and Drug Administration — Cục Quản lý Thực phẩm và Dược phẩm Mỹ'),
    ('GAD', 'Generalized Anxiety Disorder — Rối loạn Lo âu Tổng quát'),
    ('CGI-I', 'Clinical Global Impression — Improvement'),
    ('ERN', 'Error-Related Negativity — Điện thế Tiêu cực Liên quan Lỗi'),
    ('fMRI', 'Functional Magnetic Resonance Imaging'),
    ('BDNF', 'Brain-Derived Neurotrophic Factor'),
    ('SUCRA', 'Surface Under the Cumulative Ranking'),
    ('NMA', 'Network Meta-Analysis'),
])

# PHẢN BIỆN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'AJP Q1 IF ≈ 18 — cực uy tín. Tác giả NIMH/NIH — chuyên gia hàng đầu thế giới về lo âu trẻ em (Daniel Pine: >500 bài NC, h-index >100).',
    'Tổng quan TOÀN DIỆN: CBT + dược lý + kết hợp + hướng tương lai — bao quát nhất trong Đề tài về điều trị. BMC NMA 2025 (QT29) tập trung so sánh can thiệp; bài này thêm CƠ CHẾ thần kinh.',
    'CAMS: RCT lớn nhất, quan trọng nhất cho lo âu trẻ em. 80,7% kết hợp — con số then chốt cho đề cương.',
    'Cảnh báo FDA + phân tích NNT/NNH — cân bằng lợi ích/rủi ro SSRIs rất rõ ràng.',
    'Hướng tương lai (ABMT, DCS, VR, biomarker) — mở rộng kiến thức về những gì CHƯA ĐỦ bằng chứng.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'Tổng quan NARRATIVE — không có tìm kiếm hệ thống, không PRISMA. BMC NMA 2025 (QT29) chặt chẽ hơn về phương pháp.',
    'Chủ yếu bằng chứng PHƯƠNG TÂY (Mỹ, Âu, Úc) — RẤT ÍT đề cập LMIC/châu Á. VN: 0 RCT CBT cho lo âu trẻ em (Gap #1 cross-study). Zhameden 2025 (QT03): 0 RCT từ VN/ĐNA.',
    'Tập trung 3 loại lo âu (GAD, chia ly, xã hội) — không bao gồm phobia cụ thể, PTSD. Jefferies 2020 (QT35): lo âu XÃ HỘI ở VN = 30,7% — cần NC riêng.',
    'ABMT, DCS, VR — hướng tương lai nhưng chưa có ĐỦ bằng chứng, đặc biệt ở trẻ em. DCS chỉ 2 RCT nhỏ. ABMT hiệu ứng nhỏ.',
    'Không đề cập ACT — BMC NMA 2025 (QT29): ACT xếp hạng 1 (SUCRA 0,69), TRÊN CBT. Thiếu sót quan trọng.',
    'Không đề cập PE — BMC NMA 2025 (QT29): PE hiệu quả (SUCRA 0,51). Dễ triển khai tại trường VN (chi phí thấp, không cần chuyên gia).',
    'Không đề cập can thiệp TOÀN GIA ĐÌNH — JAMA 2024 (QT33): RCT toàn gia đình Cohen d = 0,53. Phù hợp văn hóa VN.',
]:
    add_red(doc, f'• {s}')

add_red(doc, 'Khoảng trống nghiên cứu / Research Gap:', bold=True)
for s in [
    'VN CHƯA CÓ RCT CBT cho lo âu trẻ em — GAP #1 cross-study, CẤP THIẾT NHẤT. Cần RCT tại trường VN: CBT nhóm (BMC NMA QT29: CBT SUCRA 0,66) + PE (SUCRA 0,51).',
    'Đánh giá hiệu quả CBT trong BỐI CẢNH VN — văn hóa, ngôn ngữ, nguồn lực. Zhameden 2025 (QT03): can thiệp trường LMIC 3/4 hiệu quả trầm cảm nhưng chỉ 1/4 lo âu.',
    'Kết hợp CBT+SSRI CHƯA được thử nghiệm tại VN — CAMS 80,7% ở Mỹ, có thể khác ở VN (tiếp cận thuốc, nhân lực tâm lý).',
    'iCBT cho VN — ÍT nhân lực tâm lý, internet phổ biến → tiềm năng lớn. Cần RCT iCBT tiếng Việt cho VTN.',
    'ACT tại VN — BMC NMA (QT29) hạng 1 nhưng MỚI ở VN. Cần đào tạo chuyên gia trước khi thử nghiệm. Thảo Vi 2025 (VN19): lạc quan (β = −0,24 gián tiếp) — gần ACT (chấp nhận + giá trị).',
    'PE (hoạt động thể chất) tại trường VN — DỄ TRIỂN KHAI NHẤT, CHI PHÍ THẤP NHẤT. Cần RCT: thể dục + yoga/mindfulness → giảm lo âu? BMC NMA (QT29): PE SUCRA 0,51. Zhu 2025 (QT05): hoạt động ngoài trời bảo vệ SKTT.',
]:
    add_red(doc, f'• {s}')

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '28_AJP_PediatricAnxiety_2024.docx')
doc.save(outpath)
import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
