# -*- coding: utf-8 -*-
"""Fix tat ca loi da phat hien trong Tong quan v2 33-trang + FULL v1:
1. VN017 Lam 2022 - 4 loi (n, dia ly, ty le, bia nam>nu)
2. VN016 Bao Quyen 2025 - 2 loi (n, ty le)
3. VN018 An Giang 2025 - 3 loi (ten tac gia, n, ty le)
4. QT035 Jefferies 2020 - 2 loi (7 nuoc, cao nhat)
5. QT005 Alharbi 2019 - 2 loi (n, ty le nho)
6. QT006 Nakie 2022 - 2 loi (n, ty le)
7. "tong quat" -> "lan toa" (2 cho)
Output: ban v3 voi cac sua DO + update bao cao loi.
26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC_33 = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_TheoChuDe_33trang_FINAL_26052026.docx')
SRC_FULL = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_TheoChuDe_FULL_v1_26052026.docx')
OUT_33 = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_TheoChuDe_33trang_v3_FixSoLieu_26052026.docx')
OUT_FULL = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_TheoChuDe_FULL_v2_FixSoLieu_26052026.docx')

RED = RGBColor(192, 0, 0)

# ============================================================
# CAC DOAN CAN THAY THE (FULL TEXT, OLD -> NEW)
# ============================================================
# Old text (key fragment to find), New corrected text (will mark RED entire new text)

REPLACEMENTS = [
    # 1. Bao Quyen 2025: n=1.252, ty le 38,5% --> n=501, ty le 86,2%
    {
        'find': 'Nguyễn Ngọc Bảo Quyên và cộng sự (2025) trên 1.252 học sinh trung học phổ thông Hà Nội bằng thang DASS-21 phiên bản tiếng Việt báo cáo tỷ lệ lo âu chung 38,5%, với nữ cao hơn nam có ý nghĩa thống kê.',
        'replace': 'Nguyễn Ngọc Bảo Quyên và cộng sự (2025) trên 501 học sinh trung học phổ thông Hà Nội bằng thang DASS-21 phiên bản tiếng Việt báo cáo tỷ lệ lo âu chung 86,2% (cao nhất trong nhóm các nghiên cứu Việt Nam tổng hợp ở đây; phân tầng: nhẹ 6,6%; vừa 25,1%; nặng 17,4%; rất nặng 37,1%), với nữ cao hơn nam có ý nghĩa thống kê (p < 0,001).',
        'tag': 'VN016 Bảo Quyên 2025: fix n và tỷ lệ'
    },
    # 2. Lam 2022: 4 loi (1.024 -> 482; ban do thi -> ban nong thon; 25,8% -> 49,0%; bia nam>nu)
    {
        'find': 'Nguyễn Danh Lâm và cộng sự (2022) trên 1.024 học sinh trung học phổ thông tại Yên Định, Thanh Hóa — khu vực bán đô thị — báo cáo tỷ lệ thấp hơn 25,8% bằng DASS-21, đồng thời ghi nhận học sinh nam có tỷ lệ cao hơn nữ ở khu vực này — một phát hiện trái ngược với xu hướng chung và đáng được nghiên cứu sâu hơn.',
        'replace': 'Nguyễn Danh Lâm và cộng sự (2022) trên 482 học sinh trung học phổ thông (2 trường) tại huyện Yên Định, Thanh Hóa — huyện bán nông thôn — báo cáo tỷ lệ lo âu chung 49,0% bằng DASS-21 (nhẹ 11,2%; vừa 25,1%; nặng 8,1%; rất nặng 4,6%). Nghiên cứu này không phân tích lo âu theo giới tính.',
        'tag': 'VN017 Lâm 2022: fix n + địa lý + tỷ lệ + bỏ claim giới'
    },
    # 3. An Giang 2025: Le Minh T -> Nguyen Dang Khoa; n=600 -> 366; 33,2% -> 61,2%
    {
        'find': 'Lê Minh T. và cộng sự (2025) trên 600 học sinh phổ thông Long Bình, An Giang bằng DASS-21 báo cáo tỷ lệ 33,2%.',
        'replace': 'Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh (2025) trên 366 học sinh THCS&THPT Long Bình, An Giang bằng DASS-21 báo cáo tỷ lệ lo âu 61,2% (nhẹ 9,3%; vừa 24,0%; nặng 12,6%; rất nặng 15,3%).',
        'tag': 'VN018 An Giang 2025: fix tác giả + n + tỷ lệ'
    },
    # 4. Jefferies 2020: 7 nuoc sai + Thai Lan cao nhat sai
    {
        'find': 'Sử dụng thang SIAS (Social Interaction Anxiety Scale) trên thanh niên ở Việt Nam, Thái Lan, Indonesia, Trung Quốc, Hoa Kỳ, Canada và Anh, các tác giả báo cáo tỷ lệ trung bình 36% có lo âu xã hội cho cả bảy quốc gia, trong đó Việt Nam đạt 30,7% — gần với mức trung bình. Thái Lan có tỷ lệ cao nhất 41%, Indonesia thấp nhất 22,9%.',
        'replace': 'Sử dụng thang SIAS (Social Interaction Anxiety Scale) trên thanh niên ở bảy quốc gia (Brazil, Trung Quốc, Indonesia, Nga, Thái Lan, Hoa Kỳ và Việt Nam), các tác giả báo cáo tỷ lệ trung bình 36,2% có lo âu xã hội, trong đó Việt Nam đạt 30,7% — gần với mức trung bình. Hoa Kỳ có tỷ lệ cao nhất (57,6%); Brazil 42,4%; Thái Lan 41,4%; Trung Quốc 32,1%; Việt Nam 30,7%; Nga 27,0%; Indonesia thấp nhất 22,9%.',
        'tag': 'QT035 Jefferies 2020: fix danh sách nước + cao nhất'
    },
    # 5. Alharbi 2019: n=366 -> 1.245; 63,1% -> 63,4%
    {
        'find': 'Alharbi và cộng sự (2019) trên 366 học sinh trung học phổ thông tại vùng Qassim, Ả Rập Saudi bằng GAD-7 báo cáo tỷ lệ 63,1%',
        'replace': 'Alharbi và cộng sự (2019) trên 1.245 học sinh trung học phổ thông tại vùng Qassim, Ả Rập Saudi bằng GAD-7 báo cáo tỷ lệ 63,4% (nhẹ 34,1%; vừa 19,5%; nặng 9,8%)',
        'tag': 'QT005 Alharbi 2019: fix n + tỷ lệ'
    },
    # 6. Nakie 2022: 643 -> 849; 38,9% -> 66,7%
    {
        'find': 'Nakie và cộng sự (2022) trên 643 học sinh trung học phổ thông tại Ethiopia bằng DASS-21 báo cáo tỷ lệ lo âu 38,9%',
        'replace': 'Nakie và cộng sự (2022) trên 849 học sinh trung học phổ thông tại Northwest Ethiopia bằng DASS-21 báo cáo tỷ lệ lo âu 66,7% (cùng với trầm cảm 41,4% và stress 52,2%)',
        'tag': 'QT006 Nakie 2022: fix n + tỷ lệ'
    },
    # 7. "lo âu tổng quát" -> "lo âu lan tỏa" (2 chỗ)
    {
        'find': 'lo âu tổng quát (chủ yếu qua GAD-7',
        'replace': 'lo âu lan tỏa (chủ yếu qua GAD-7',
        'tag': 'Thuật ngữ: tổng quát → lan tỏa (chỗ 1)'
    },
    {
        'find': 'lo âu (lo âu tổng quát 56%, lo âu xã hội 41%)',
        'replace': 'lo âu (lo âu lan tỏa 56%, lo âu xã hội 41%)',
        'tag': 'Thuật ngữ: tổng quát → lan tỏa (chỗ 2)'
    },
]

def apply_replacements_in_paragraph(p, replacements, applied_tracker):
    """Apply MULTIPLE replacements in a single paragraph.
    Approach: build a list of segments [(text, color_red_or_not)] from original p.text,
    then rebuild paragraph runs from segments."""
    full = p.text
    # Find all matches with positions
    matches = []
    for rep in replacements:
        pos = full.find(rep['find'])
        while pos != -1:
            matches.append((pos, pos + len(rep['find']), rep))
            pos = full.find(rep['find'], pos + 1)
    if not matches:
        return
    # Sort matches by position
    matches.sort(key=lambda x: x[0])
    # Check no overlap
    safe_matches = []
    last_end = -1
    for m in matches:
        if m[0] >= last_end:
            safe_matches.append(m)
            last_end = m[1]
    if not safe_matches:
        return
    # Build segments
    segments = []  # (text, is_red)
    cursor = 0
    for start, end, rep in safe_matches:
        if cursor < start:
            segments.append((full[cursor:start], False))
        segments.append((rep['replace'], True))
        applied_tracker.append(rep['tag'])
        cursor = end
    if cursor < len(full):
        segments.append((full[cursor:], False))
    # Clear all runs
    for r in p.runs:
        r.text = ''
    # Add segments as new runs
    for text, is_red in segments:
        if not text:
            continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True

def apply_replacements(src_path, out_path, replacements):
    doc = Document(src_path)
    applied = []
    for p in doc.paragraphs:
        apply_replacements_in_paragraph(p, replacements, applied)
    doc.save(out_path)
    return applied

# Run for both files
print("=== Tong quan v2 33-trang ===")
applied_33 = apply_replacements(SRC_33, OUT_33, REPLACEMENTS)
for a in applied_33:
    print(f"  ✓ {a}")
print(f"Total: {len(applied_33)} fixes")
print(f"Saved: {OUT_33}")
print()

print("=== Tong quan FULL v1 ===")
applied_full = apply_replacements(SRC_FULL, OUT_FULL, REPLACEMENTS)
for a in applied_full:
    print(f"  ✓ {a}")
print(f"Total: {len(applied_full)} fixes")
print(f"Saved: {OUT_FULL}")
