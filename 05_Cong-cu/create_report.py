# -*- coding: utf-8 -*-
"""Create comprehensive comparison report WITH Vietnamese diacritics"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

# TITLE
doc.add_heading('BÁO CÁO TỔNG HỢP', level=0)
doc.add_heading('Phân tích so sánh 11 nghiên cứu về lo âu và trầm cảm ở học sinh', level=1)
doc.add_paragraph()

# PART 1
doc.add_heading('PHẦN 1: BẢNG SO SÁNH TỔNG HỢP', level=1)
doc.add_paragraph('Bảng dưới đây tổng hợp tỷ lệ lo âu, trầm cảm, cỡ mẫu, công cụ đo lường, và phát hiện nổi bật của 11 nghiên cứu.')

data = [
    ['STT', 'Tác giả', 'Quốc gia', 'N', 'Công cụ', 'Lo âu %', 'Trầm cảm %', 'Căng thẳng %', 'Giới > lo âu'],
    ['1', 'Jenkins 2023', 'Hoa Kỳ', '75', 'PHQ-9A GAD-10', '50,6', '44,0', '-', 'Nữ'],
    ['2', 'Saikia 2023', 'Ấn Độ', '360', 'DASS-21', '24,4', '22,2', '6,9', 'NAM*'],
    ['3', 'Mandaknalli 2021', 'Ấn Độ', '450', 'GAD-7', '~100', '-', '-', 'Nữ'],
    ['4', 'NSCH 2020', 'Hoa Kỳ', '55.162', 'Khảo sát QG', '16,1', '8,4', '-', 'Nữ'],
    ['5', 'Alharbi 2019', 'Ả Rập Saudi', '1.245', 'PHQ-9 GAD-7', '63,5', '74,0', '-', 'Nữ'],
    ['6', 'Nakie 2022', 'Ethiopia', '849', 'DASS-21', '66,7', '41,4', '52,2', 'Nữ'],
    ['7', 'Chen 2023', 'Trung Quốc', '63.205', 'PHQ-9 GAD-7', '13,9', '23,0', '-', 'Nữ'],
    ['8', 'Wen 2020', 'Trung Quốc', '900', 'MHT+LPA', '24,78', '-', '-', 'NAM*'],
    ['9', 'Qiu 2022', 'Trung Quốc', '2.079', 'CES-D SAS', '13,4', '26,0', '-', '-'],
    ['10', 'Xu 2021', 'Trung Quốc', '373.216', 'GAD-7', '9,89', '-', '-', 'NAM*'],
    ['11', 'Bhardwaj 2020', 'Ấn Độ', '288', 'DASS-21', '73,3', '64,9', '74,7', '-'],
]

table = doc.add_table(rows=len(data), cols=len(data[0]))
table.style = 'Table Grid'
for i, row_data in enumerate(data):
    for j, val in enumerate(row_data):
        cell = table.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True
        if 'NAM' in val and i > 0:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('NAM* = Nam có tỷ lệ lo âu cao hơn nữ — trái ngược xu hướng toàn cầu (Saikia, Wen, Xu)')
run.font.size = Pt(9)
run.font.color.rgb = RGBColor(0xCC, 0, 0)
run.bold = True

p = doc.add_paragraph()
run = p.add_run('Nhận xét: Tỷ lệ lo âu dao động rộng từ 9,89% (Xu — cỡ mẫu lớn nhất) đến 73,3% (Bhardwaj — trường công lập). Sự khác biệt lớn do: (1) công cụ đo khác nhau; (2) ngưỡng cắt khác nhau; (3) bối cảnh văn hóa-xã hội; (4) cỡ mẫu.')
run.font.size = Pt(10)

doc.add_page_break()

# PART 2: CHARTS
doc.add_heading('PHẦN 2: BIỂU ĐỒ SO SÁNH', level=1)

charts = [
    ('Charts/01_anxiety_comparison.png',
     'Biểu đồ 1: Tỷ lệ lo âu — 11 nghiên cứu. Xanh <20%, Cam 20-50%, Đỏ >50%. Đường xanh lá = trung bình toàn cầu ~25% (Racine et al., JAMA Pediatrics, 2021).'),
    ('Charts/02_depression_comparison.png',
     'Biểu đồ 2: Tỷ lệ trầm cảm. Alharbi (74%) và Bhardwaj (64,9%) cao bất thường — có thể do ngưỡng PHQ-9 ≥5 bao gồm cả triệu chứng nhẹ.'),
    ('Charts/04_regional_comparison.png',
     'Biểu đồ 3: Tỷ lệ lo âu trung bình theo vùng địa lý. Ấn Độ và Ethiopia có tỷ lệ cao nhất. Trung Quốc thấp nhất nhờ cỡ mẫu lớn và công cụ chuẩn.'),
    ('Charts/03_sample_vs_anxiety.png',
     'Biểu đồ 4: Mối quan hệ cỡ mẫu — tỷ lệ lo âu. Xu hướng: cỡ mẫu càng lớn, tỷ lệ báo cáo càng thấp. Nghiên cứu nhỏ có thể phóng đại tỷ lệ.'),
]

for img_file, caption in charts:
    img_path = os.path.join(BASE, img_file)
    if os.path.exists(img_path):
        doc.add_picture(img_path, width=Inches(6.0))
        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        p = doc.add_paragraph()
        run = p.add_run(caption)
        run.font.size = Pt(9)
        run.font.italic = True
        doc.add_paragraph()

doc.add_page_break()

# PART 3: RESEARCH PROPOSAL
doc.add_heading('PHẦN 3: ĐỀ CƯƠNG NGHIÊN CỨU ĐỀ XUẤT', level=1)

doc.add_heading('3.1. Khoảng trống nghiên cứu được xác định', level=2)
gaps = [
    'Thiếu nghiên cứu dọc (longitudinal) — tất cả 11 bài đều là cắt ngang, không thể xác lập nhân quả.',
    'Thiếu nghiên cứu can thiệp (intervention) — không bài nào đánh giá hiệu quả chương trình can thiệp.',
    'Thiếu dữ liệu từ Đông Nam Á — không có bài nào từ Việt Nam, Thái Lan, Philippines.',
    'Công cụ đo lường không thống nhất — PHQ-9, GAD-7, DASS-21, CES-D, SAS được sử dụng khác nhau, gây khó so sánh.',
    'Thiếu phân tích đa tầng (multilevel) — hầu hết chưa xem xét hiệu ứng cụm (trường học).',
    'Mối quan hệ giới tính — lo âu phức tạp hơn giả định — 3/11 bài cho thấy nam > nữ, cần nghiên cứu thêm.',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('3.2. Đề cương nghiên cứu đề xuất', level=2)

items = [
    ('Tên đề tài:', 'Tỷ lệ hiện mắc và các yếu tố liên quan của lo âu, trầm cảm ở học sinh trung học cơ sở tại [địa phương], Việt Nam: Nghiên cứu cắt ngang tại trường học'),
    ('Mục tiêu:', '(1) Xác định tỷ lệ lo âu và trầm cảm; (2) Xác định các yếu tố nguy cơ và bảo vệ; (3) So sánh giữa các nhóm giới tính, vùng miền, điều kiện kinh tế xã hội'),
    ('Thiết kế:', 'Cắt ngang, tại trường học, lấy mẫu cụm nhiều giai đoạn'),
    ('Cỡ mẫu:', 'Tối thiểu 1.000 học sinh (tham khảo Saikia n=360, Alharbi n=1.245, Nakie n=849)'),
    ('Công cụ:', 'PHQ-9A + GAD-7 (để so sánh với Alharbi, Chen, Xu) HOẶC DASS-21 (để so sánh với Saikia, Nakie, Bhardwaj). Cần xác thực phiên bản tiếng Việt.'),
    ('Phân tích:', 'Hồi quy logistic đa biến (tham khảo Nakie, Chen). Xem xét mô hình đa tầng (multilevel) để kiểm soát hiệu ứng cụm trường học.'),
    ('Biến số:', 'Giới tính, tuổi, thành thị/nông thôn, cấu trúc gia đình, sử dụng internet/game (Chen), chất lượng giấc ngủ (Chen), bắt nạt (Chen), phong cách nuôi dạy con (Qiu), khả năng phục hồi (Qiu)'),
    ('Đạo đức:', 'Phê duyệt IRB, đồng ý của phụ huynh và học sinh, bảo mật dữ liệu'),
    ('Hạn chế dự kiến:', 'Thiết kế cắt ngang (không xác lập nhân quả), tự báo cáo (không phải chẩn đoán lâm sàng), một địa phương (hạn chế khái quát hóa)'),
    ('Ý nghĩa:', 'Cung cấp dữ liệu cơ sở cho chính sách sức khỏe tâm thần học đường tại Việt Nam. Đóng góp vào y văn khu vực Đông Nam Á — hiện còn rất thiếu dữ liệu.'),
]
for title, content in items:
    p = doc.add_paragraph()
    run = p.add_run(title + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(content)
    run.font.size = Pt(11)

doc.add_heading('3.3. Khung thời gian đề xuất', level=2)
tl = [
    ['Giai đoạn', 'Thời gian', 'Nội dung'],
    ['1. Chuẩn bị', 'Tháng 1-2', 'Xác thực công cụ, phê duyệt đạo đức, tập huấn'],
    ['2. Thu thập', 'Tháng 3-5', 'Khảo sát tại trường học'],
    ['3. Phân tích', 'Tháng 6-7', 'Nhập liệu, làm sạch, phân tích thống kê'],
    ['4. Báo cáo', 'Tháng 8-9', 'Viết báo cáo, gửi tạp chí'],
]
t2 = doc.add_table(rows=len(tl), cols=3)
t2.style = 'Table Grid'
for i, row in enumerate(tl):
    for j, val in enumerate(row):
        cell = t2.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if i == 0:
            run.bold = True

out = os.path.join(BASE, 'DocFiles', 'BAO_CAO_TONG_HOP_SO_SANH.docx')
doc.save(out)
print(f'DONE: {out}')
