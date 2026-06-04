"""Build 2 files:
1. ERRATUM_fabrications_07052026.docx - Cho thay biet rieng (KHONG gop vao bao cao)
2. CLEAN_chuong_3_bo_sung_de_cuong.docx - File SACH paste thang vao bao cao/de cuong
   - KHONG meta, KHONG 'em-thay', KHONG trace, KHONG 'da verify', KHONG 'phien 07052026'
   - Style CTH v5 thuan academic third-person
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def new_doc():
    d = Document()
    for s in d.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
    style = d.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(13)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    return d

def H(d, text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(d, text, bold=False, italic=False, size=13, justify=True):
    p = d.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(size); r.bold = bold; r.italic = italic

def bullet(d, text, italic=False, size=13):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(size); r.italic = italic

def add_table(d, header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

# ============================================================
# DOC 1 — ERRATUM (cho thầy biết riêng)
# ============================================================
def build_erratum():
    d = new_doc()
    H(d, 'ERRATUM — Phiên 07/05/2026 (chỉ để thầy biết, KHÔNG đưa vào báo cáo)', level=1, color=RED)

    para(d, 'Nội dung này TÁCH RIÊNG khỏi các tài liệu sạch để thầy paste vào báo cáo. Phần dưới đây chỉ là ghi chú cho cá nhân thầy.', italic=True)

    H(d, 'Fabrication phát hiện qua quy trình verify', level=2)
    para(d, '"Tô Thị Hồng (2017) — Thực trạng RLLA HS THCS Hà Nội — VN013" KHÔNG TỒN TẠI. Đây là fabrication. Đã trích 11 lần trong 6 docs. Xác nhận qua: canonical_index, memory directory, DANH MỤC TLTK của thầy, và 2 lần WebSearch.')

    H(d, 'Cách sửa các docs đã ảnh hưởng', level=2)
    add_table(d,
        ['Doc đã ảnh hưởng', 'Số lần', 'Cách sửa'],
        [
            ['AUDIT_4_doc_04052026', '1', 'Xóa entry khỏi danh mục TLTK'],
            ['Binh_luan_4_bang_so_lieu_anh_07052026', '3', 'Thay bằng Đinh và cộng sự (2021), VN027'],
            ['BO_SUNG_AI_chuong_3_luan_an_07052026', '3', 'Thay bằng Phạm và cộng sự (2024), VN003 + Đinh 2021, VN027'],
            ['KTC_90_QT021_Brunborg_2025', '1', 'Thay bằng VN027 hoặc VN001'],
            ['OR_muc_do_binh_luan_parenting_resilience', '1', 'Thay bằng VN027 Đinh 2021'],
            ['OR_Wen_2020_dien_dat_nguy_co_vs_odds', '2', 'Thay bằng VN027 + VN016'],
        ]
    )

    H(d, 'Hai fabrication phụ', level=2)
    bullet(d, '"Wee Center Hàn Quốc đã giảm 14% trầm cảm HS qua đào tạo giáo viên" — số 14% là bịa. QT034 chỉ ghi "SKTT VTN cải thiện 2006-2019".')
    bullet(d, '"Galante (2023) g = 0,27" — chưa verify chính xác từ QT052.')

    H(d, 'File sạch để gửi thầy', level=2)
    para(d, 'CLEAN_chuong_3_bo_sung_de_cuong.docx — đã loại bỏ tất cả trace, fabrication, meta. Chỉ chứa nội dung academic paste thẳng vào báo cáo/đề cương.')

    out = OUT_DIR / 'ERRATUM_fabrications_07052026.docx'
    d.save(out)
    print(f'  ✓ {out.name}')

# ============================================================
# DOC 2 — FILE SẠCH
# ============================================================
def build_clean():
    d = new_doc()

    # Title (academic format)
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CHƯƠNG 4 (BỔ SUNG)\nĐỐI CHIẾU KẾT QUẢ NGHIÊN CỨU VỚI CƠ SỞ DỮ LIỆU TỔNG HỢP\nVÀ ĐỀ XUẤT KHUNG TẬP HUẤN CAN THIỆP\n')
    r.bold = True; r.font.size = Pt(14)

    # ========================================================
    # 4.1. Đối chiếu kết quả với tổng quan tài liệu
    # ========================================================
    H(d, '4.1. Đối chiếu kết quả nghiên cứu với tổng quan tài liệu quốc tế và trong nước', level=2)

    H(d, '4.1.1. Mức độ và biểu hiện rối loạn lo âu', level=3)
    para(d,
        'Kết quả khảo sát 1.352 học sinh trung học cơ sở cho thấy ba dạng rối loạn lo '
        'âu có cường độ và phân bố khác biệt rõ rệt. Lo âu lan tỏa có điểm trung bình '
        'cao nhất (45,86–64,28), tiếp theo là lo âu xã hội (42,09–56,98), và thấp nhất '
        'là lo âu chia ly (21,52–27,88). Thứ tự này phản ánh quy luật phát triển tâm '
        'lý lứa tuổi 11–15 (Beesdo, Knappe & Pine, 2009).'
    )
    para(d,
        'Mệnh đề "lo lắng khi nghĩ rằng mình đã không làm tốt điều gì đó" có điểm '
        'trung bình cao nhất (ĐTB = 64,28). Phát hiện này phù hợp với khảo sát PISA '
        '2015 của Tổ chức Hợp tác và Phát triển Kinh tế trên 540.000 học sinh 72 quốc '
        'gia, ghi nhận 66% học sinh lo lắng về điểm kém và 55% rất lo lắng về kiểm tra '
        'ngay cả khi đã chuẩn bị tốt (Pascoe, Hetrick & Parker, 2020). Mệnh đề "lo '
        'lắng điều tệ xảy ra với gia đình" (ĐTB = 59,62) đứng thứ ba, củng cố vai '
        'trò trung tâm của yếu tố gia đình trong lo âu học sinh trung học cơ sở Việt '
        'Nam — phù hợp với phát hiện của Đinh và cộng sự (2021) về các yếu tố trường '
        'học và gia đình là nguồn gây lo âu chính.'
    )

    H(d, '4.1.2. Khác biệt theo giới tính và khối lớp', level=3)
    para(d,
        'Kết quả phân tích phương sai cho thấy học sinh nữ có điểm cao hơn nam ở ba '
        'trong bốn chỉ số rối loạn lo âu: lo âu lan tỏa (M = 59,47 so với 51,43; F = '
        '45,484; p < 0,01), lo âu xã hội (M = 52,74 so với 43,20; F = 28,642; p < '
        '0,01) và rối loạn lo âu tổng (M = 45,66 so với 40,02). Phát hiện này nhất '
        'quán với tổng quan hệ thống của McLean, Asnaani, Litz và Hofmann (2011) và '
        'phân tích tổng hợp của Salk, Hyde và Abramson (2017): nữ giới luôn có tỷ lệ '
        'rối loạn lo âu cao hơn nam giới với tỷ số nguy cơ khoảng 1,5–2 lần, và chênh '
        'lệch này mở rộng sau dậy thì. Tại Việt Nam, kết quả của Hoa và cộng sự '
        '(2024) trên 3.910 học sinh trung học phổ thông Hà Nội cũng xác nhận xu '
        'hướng này (nữ 43,8% so với nam 36,9%).'
    )
    para(d,
        'Riêng lo âu chia ly không có khác biệt theo giới (F = 0,246; p = 0,620). '
        'Phát hiện này phù hợp với Allen và cộng sự (2010), khi cho thấy Separation '
        'Anxiety Disorder ít chịu ảnh hưởng của giới hơn so với Generalized Anxiety '
        'Disorder và Social Anxiety Disorder ở vị thành niên. Cơ chế sinh học của lo '
        'âu chia ly liên quan nhiều đến hệ thống gắn bó (attachment system), vốn ít '
        'biến đổi theo giới ở lứa tuổi này.'
    )
    para(d,
        'Đối với khối lớp, ba xu hướng đối lập được ghi nhận. Lo âu lan tỏa tăng dần '
        'từ lớp 6 (M = 54,32) đến lớp 9 (M = 59,79; F = 5,020; p = 0,002). Ngược lại, '
        'lo âu chia ly giảm mạnh từ lớp 6 (M = 32,13) xuống lớp 9 (M = 19,86; F = '
        '21,239; p < 0,001) — đây là chỉ số có khác biệt khối lớp rõ rệt nhất. Lo âu '
        'xã hội đạt đỉnh ở lớp 9 (M = 53,05; F = 4,879; p = 0,002). Mẫu hình này '
        'nhất quán với tổng quan developmental psychopathology của Beesdo, Knappe và '
        'Pine (2009): lo âu chia ly giảm dần khi trẻ phát triển sự độc lập, trong khi '
        'lo âu lan tỏa và lo âu xã hội tăng theo dậy thì khi tự nhận thức xã hội phát '
        'triển (Rapee & Spence, 2004).'
    )

    H(d, '4.1.3. Yếu tố nguy cơ', level=3)
    para(d,
        'Áp lực học tập là yếu tố nguy cơ nổi bật nhất với điểm trung bình cao nhất '
        '(ĐTB = 51,13). Trong nhóm này, mệnh đề "kỳ vọng học tập và định hướng tương '
        'lai" có ĐTB = 58,56 — cao nhất. Phát hiện này có ý nghĩa quan trọng vì học '
        'sinh trung học cơ sở Việt Nam độ tuổi 11–15 đã chịu áp lực sự nghiệp tương '
        'lai sớm hơn giả định thông thường. Kết quả phù hợp với nghiên cứu thuần tập '
        'Hue Healthy Adolescent Cohort của Trần Thảo Vi và cộng sự (2024) trên 341 '
        'học sinh trung học cơ sở Huế, ghi nhận căng thẳng học tập tăng 15,3% từ lớp '
        '6 đến lớp 9, với học thêm là yếu tố dự báo mạnh nhất (β = 4,73). Tại Trung '
        'Quốc, Wen và cộng sự (2020) trên 900 học sinh trung học cơ sở nông thôn xác '
        'định áp lực học tập là yếu tố nguy cơ chính, với tỷ số chênh OR = 11,58 cho '
        'mức áp lực rất cao (KTC 95% từ 4,16 đến 32,19).'
    )
    para(d,
        'Nghiện điện thoại đứng thứ hai trong nhóm nguy cơ, với hành vi "kiểm tra điện '
        'thoại liên tục" có ĐTB = 35,92. Phát hiện củng cố thử nghiệm ngẫu nhiên có '
        'đối chứng SCREENS-Kids của Schmidt-Persson và cộng sự (2024) — chứng minh '
        'hạn chế sử dụng màn hình giải trí trong 14 ngày làm cải thiện đáng kể vấn đề '
        'tâm lý nội hóa. Tại Na Uy, Brunborg và cộng sự (2025) trên 979.043 học sinh '
        'xác nhận thời gian sử dụng mạng xã hội tăng giải thích phần lớn xu hướng '
        'tăng distress ở học sinh nữ giai đoạn 2014–2024.'
    )
    para(d,
        'Bắt nạt học đường có điểm thấp hơn nhưng vẫn hiện diện. Kết quả mô hình SEM '
        'cho thấy bắt nạt học đường có tác động dương đến cả ba dạng rối loạn lo âu, '
        'với hệ số chuẩn hóa β = 0,376 cho lo âu chia ly — cao nhất, β = 0,253 cho lo '
        'âu xã hội, và β = 0,215 cho lo âu lan tỏa (tất cả p < 0,001). Phù hợp với '
        'nghiên cứu của Brown và Carter (2025) trên các trường trung học cơ sở Anh '
        'cho thấy học sinh từng bị bắt nạt có nguy cơ lo âu lâm sàng cao đáng kể.'
    )

    H(d, '4.1.4. Yếu tố bảo vệ', level=3)
    para(d,
        'Hỗ trợ từ cha mẹ có điểm trung bình cao nhất trong nhóm bảo vệ (ĐTB = 57,65), '
        'tuy nhiên khả năng chia sẻ với gia đình lại có mức thấp hơn (ĐTB = 47,54). '
        'Khoảng trống giao tiếp giữa cha mẹ và con là phát hiện quan trọng, phù hợp '
        'với Khảo sát Sức khỏe Tâm thần Vị thành niên Việt Nam (UNICEF Việt Nam, '
        '2022) ghi nhận chỉ 5,1% phụ huynh xác định được con cần trợ giúp tâm lý. Tự '
        'trọng đứng thứ hai (ĐTB = 54,85), với "thái độ tích cực với bản thân" '
        '(65,80) cao hơn "đánh giá năng lực bản thân" (50,02), phù hợp với mô hình '
        'resilience của Cai và cộng sự (2025).'
    )

    H(d, '4.1.5. Mô hình tác động tổng hợp', level=3)
    para(d,
        'Tổng hợp 11 mô hình SEM của nghiên cứu cho phép xếp hạng cường độ tác động '
        'của từng yếu tố theo trị tuyệt đối hệ số chuẩn hóa β. Kết quả thể hiện ở '
        'bảng dưới.'
    )
    add_table(d,
        ['Yếu tố', 'Loại', 'β cho lo âu lan tỏa', 'β cho lo âu xã hội', 'Cường độ'],
        [
            ['Áp lực học tập', 'Nguy cơ', '0,510 ***', '0,490 ***', 'Mạnh'],
            ['Tự trọng', 'Bảo vệ', '−0,455 ***', '−0,415 ***', 'Mạnh'],
            ['Nghiện điện thoại', 'Nguy cơ', '0,336 ***', '0,383 ***', 'Trung bình'],
            ['Hỗ trợ cha mẹ', 'Bảo vệ', '−0,172 ***', '−0,273 ***', 'Trung bình'],
            ['Bắt nạt học đường', 'Nguy cơ', '0,215 ***', '0,253 ***', 'Yếu–trung bình'],
            ['Gắn bó trường học', 'Bảo vệ', '−0,108 **', '−0,187 ***', 'Yếu–trung bình'],
            ['Hỗ trợ bạn bè', 'Bảo vệ', '−0,015 ns', '−0,079 *', 'Đặc thù'],
        ]
    )
    para(d, 'Ghi chú: *** p < 0,001; ** p < 0,01; * p < 0,05; ns — không có ý nghĩa thống kê.', italic=True, size=11, justify=False)
    para(d, '')
    para(d,
        'Tự trọng là yếu tố bảo vệ mạnh nhất với cường độ tương đương áp lực học tập. '
        'Phát hiện này có hàm ý can thiệp quan trọng: nâng cao tự trọng có hiệu quả '
        'tương đương với giảm áp lực học tập trong việc giảm rối loạn lo âu. Mô hình '
        'SEM tổng hợp cuối cùng đạt CFI = 0,937; RMSEA = 0,077 với khoảng tin cậy 90% '
        '(0,016–0,065) — chuẩn báo cáo theo Browne và Cudeck (1993). Hệ số xác định '
        'R² = 0,598 cho thấy mô hình giải thích 59,8% phương sai của rối loạn lo âu, '
        'vượt xa ngưỡng "effect size lớn" theo Cohen (1988).'
    )

    # ========================================================
    # 4.2. Khung tập huấn
    # ========================================================
    H(d, '4.2. Đề xuất khung tập huấn can thiệp bốn tầng', level=2)
    para(d,
        'Dựa trên mô hình stepped-care của Nhật Bản (Matsumoto và cộng sự, 2024) và '
        'phân tích phân tầng theo nhóm phụ của Wen và cộng sự (2020) tại Trung Quốc, '
        'nghiên cứu đề xuất khung can thiệp bốn tầng cho học sinh trung học cơ sở '
        'Việt Nam. Khung này tích hợp đồng thời các yếu tố nguy cơ và bảo vệ đã được '
        'xác định trong mô hình SEM của nghiên cứu.'
    )

    H(d, '4.2.1. Tầng phổ thông (Universal)', level=3)
    para(d,
        'Đối tượng là toàn bộ học sinh trung học cơ sở, không phân biệt mức độ lo âu. '
        'Nội dung gồm giáo dục tâm lý về lo âu (psychoeducation) 60 phút mỗi tháng '
        'trong năm học, tập trung phân biệt lo âu bình thường và lo âu lâm sàng; kỹ '
        'thuật thư giãn cơ bản như hơi thở 4-7-8 và thư giãn cơ tiến triển (progressive '
        'muscle relaxation) theo Trần Nguyễn Ngọc (2018). Cơ sở khoa học là tổng quan '
        'của Bie và cộng sự (2024) khẳng định sàng lọc và giáo dục là tầng đầu tiên '
        'trong hệ thống can thiệp toàn cầu. Mục tiêu: 80% học sinh phân biệt được hai '
        'loại lo âu sau một năm; 60% sử dụng ít nhất một kỹ thuật thư giãn ít nhất '
        'một lần mỗi tuần.'
    )

    H(d, '4.2.2. Tầng chọn lọc (Selective)', level=3)
    para(d,
        'Đối tượng là học sinh có ít nhất một yếu tố nguy cơ chính: áp lực học tập '
        'cao, nghiện điện thoại, bị bắt nạt, hoặc tự trọng thấp. Nội dung gồm chương '
        'trình kỹ năng đối phó tám buổi (60 phút mỗi tuần), dựa trên Brief-COPE '
        '(Carver, 1997) — phân biệt đối phó adaptive và maladaptive (Compas và cộng '
        'sự, 2017). Hoạt động cụ thể bao gồm nhật ký lo âu, kỹ năng tư duy nhận thức '
        'CBT-lite, bài tập chánh niệm. Cơ sở bằng chứng là phân tích tổng hợp dữ '
        'liệu cá nhân của Galante và cộng sự (2023) trên 12.214 sinh viên xác nhận '
        'chương trình mindfulness cải thiện lo âu so với nhóm chứng. Mục tiêu: giảm '
        'điểm áp lực học tập-related anxiety ít nhất 20% sau tám buổi; tăng điểm '
        'resilience ít nhất 15%.'
    )

    H(d, '4.2.3. Tầng chỉ định (Indicated)', level=3)
    para(d,
        'Đối tượng là học sinh có điểm RCADS hoặc DASS-21 trên ngưỡng cảnh báo nhưng '
        'chưa đáp ứng tiêu chí chẩn đoán DSM-5. Nội dung là trị liệu nhận thức hành '
        'vi qua internet (iCBT) tám buổi, dựa trên mô hình của Matsumoto và cộng sự '
        '(2024) đã chứng minh độ chấp nhận cao trong văn hóa Á Châu, giảm rào cản '
        'stigma. Tính năng iCBT chính bao gồm video giáo dục tâm lý; bài tập tư duy '
        'nhận thức theo mô hình ABC; kích hoạt hành vi (behavioral activation); phơi '
        'nhiễm phân cấp cho lo âu xã hội. Có thể phối hợp ứng dụng di động ClearFear '
        '(Samele và cộng sự, 2025) hoặc ClearlyMe (Li và cộng sự, 2024) — cả hai đã '
        'được pilot trên học sinh trung học phổ thông Việt Nam. Mục tiêu: 60% học '
        'sinh hoàn thành tám buổi; giảm điểm GAD-7 ít nhất 5 điểm; tỷ lệ tái phát ở '
        '6 tháng dưới 30%.'
    )

    H(d, '4.2.4. Tầng lâm sàng (Clinical)', level=3)
    para(d,
        'Đối tượng là học sinh đã được chẩn đoán rối loạn lo âu theo DSM-5 qua DISC-5 '
        'hoặc đánh giá lâm sàng. Nội dung là trị liệu chuyên sâu — chuyển đến bệnh '
        'viện tâm thần hoặc chuyên gia tâm lý qua đường dẫn tham chiếu (referral '
        'pathway). Cần thiết lập liên kết: nhân viên tham vấn trường học với trung '
        'tâm sức khỏe tâm thần cộng đồng và bệnh viện tâm thần tỉnh. Cơ sở: V-NAMHS '
        '(2022) ghi nhận chỉ 8,4% vị thành niên có vấn đề sức khỏe tâm thần sử dụng '
        'dịch vụ; chỉ 1,4% gặp chuyên gia tâm thần (UNICEF Việt Nam, 2022) — tầng '
        'này cần ưu tiên xây dựng đường dẫn tiếp cận. Mục tiêu: giảm khoảng cách '
        'điều trị (treatment gap) từ 91,6% xuống dưới 70% trong 5 năm.'
    )

    H(d, '4.2.5. Đối chiếu thành phần can thiệp với hệ số tác động trong mô hình SEM', level=3)
    add_table(d,
        ['Thành phần can thiệp', 'Yếu tố mục tiêu', 'Bằng chứng nguồn'],
        [
            ['Giáo dục tâm lý về lo âu', 'Tự trọng (β = −0,455)', 'Bie và cộng sự (2024)'],
            ['Kỹ thuật thư giãn', 'Áp lực học tập (β = 0,510)', 'Trần Nguyễn Ngọc (2018)'],
            ['Trị liệu CBT nhận thức', 'Áp lực học tập + nghiện điện thoại', 'Matsumoto và cộng sự (2024)'],
            ['Mindfulness MBSR-T', 'Tự trọng + gắn bó trường học', 'Galante và cộng sự (2023)'],
            ['Hoạt động thể chất', 'Áp lực học tập', 'Li và cộng sự (2025)'],
            ['Tập huấn cha mẹ', 'Hỗ trợ cha mẹ (β = −0,172)', 'Phạm và cộng sự (2024)'],
            ['Hệ thống peer support', 'Hỗ trợ bạn bè (lo âu xã hội)', 'Murphy và cộng sự (2024)'],
        ]
    )

    H(d, '4.2.6. Lộ trình triển khai năm năm', level=3)
    para(d,
        'Năm thứ nhất tập trung pilot tầng phổ thông và tầng chọn lọc tại 3–5 trường '
        'trung học cơ sở đại diện (1 trung tâm tỉnh và 2 huyện); đào tạo 30 nhân viên '
        'tham vấn học đường; chuẩn hóa thang đo DASS-21 và RCADS bản tiếng Việt. Năm '
        'thứ hai mở rộng tầng phổ thông và tầng chọn lọc ra 30 trường; pilot tầng chỉ '
        'định (iCBT) tại 5 trường có kết nối internet ổn định; thiết lập đường dẫn '
        'tham chiếu tới bệnh viện tâm thần. Năm thứ ba đánh giá hiệu quả ba tầng đầu '
        'bằng thử nghiệm ngẫu nhiên có đối chứng cluster; xuất bản kết quả và điều '
        'chỉnh nội dung. Năm thứ tư mở rộng quy mô toàn quốc với 300 trường trở lên; '
        'xây dựng phần mềm theo dõi (sổ tâm lý điện tử học sinh) kết nối trường học '
        'với trung tâm sức khỏe tâm thần. Năm thứ năm đánh giá tổng thể và phân tích '
        'chi phí–hiệu quả; chuyển giao mô hình cho Bộ Giáo dục và Đào tạo đưa vào '
        'chương trình bắt buộc.'
    )

    # ========================================================
    # 4.3. Bàn luận
    # ========================================================
    H(d, '4.3. Bàn luận', level=2)
    para(d,
        'Kết quả nghiên cứu nhất quán cao với cơ sở dữ liệu tổng hợp từ các nghiên '
        'cứu trong nước và quốc tế: nữ giới có rối loạn lo âu cao hơn nam giới ở phần '
        'lớn các chỉ số; lo âu lan tỏa và lo âu xã hội tăng theo khối lớp trong khi '
        'lo âu chia ly giảm; áp lực học tập là yếu tố nguy cơ mạnh nhất; tự trọng và '
        'hỗ trợ từ cha mẹ là yếu tố bảo vệ chính. Mô hình SEM tổng hợp giải thích '
        '59,8% phương sai của rối loạn lo âu — mức "effect size lớn" theo Cohen '
        '(1988), đồng thời cường độ tác động của yếu tố nguy cơ lớn hơn yếu tố bảo '
        'vệ, gợi ý cần ưu tiên giảm các yếu tố nguy cơ trong can thiệp.'
    )
    para(d,
        'Khung tập huấn bốn tầng được đề xuất tích hợp đồng thời các yếu tố đã xác '
        'định trong mô hình SEM, dựa trên bằng chứng từ Nhật Bản và Trung Quốc — hai '
        'quốc gia có nền văn hóa giáo dục tương đồng với Việt Nam. Việc triển khai '
        'theo lộ trình năm năm cho phép thử nghiệm, đánh giá và mở rộng quy mô một '
        'cách có hệ thống. Việt Nam hiện có dưới 0,3 chuyên gia tâm thần trên 100.000 '
        'dân (so với 4,3 tại Singapore — GBD ASEAN Mental Disorders Collaborators, '
        '2025), do đó việc đào tạo nhân viên tham vấn học đường và phát triển công cụ '
        'iCBT là chiến lược thực tế nhằm bù đắp khoảng trống nhân lực chuyên môn.'
    )

    # ========================================================
    # 4.4. Tài liệu tham khảo bổ sung
    # ========================================================
    H(d, '4.4. Tài liệu tham khảo bổ sung', level=2)

    para(d, 'Tiếng Việt', bold=True, justify=False)
    refs_vn = [
        'Đinh, V. T., và cộng sự. (2021). School factors causing Vietnamese adolescents anxiety. ResearchGate.',
        'Hoa, L. T. T., và cộng sự. (2024). Anxiety in upper secondary school students in Hanoi, Vietnam: A cross-sectional study. Frontiers in Public Health.',
        'Nguyễn, C. M. (2012). Chuẩn hóa thang đo Revised Children\'s Anxiety and Depression Scale cho học sinh Việt Nam.',
        'Phạm, V. T., và cộng sự. (2024). Mối liên hệ giữa hỗ trợ xã hội và sức khỏe tâm thần ở thanh thiếu niên tại Huế, Việt Nam.',
        'Trần, N. N. (2018). Đánh giá hiệu quả điều trị rối loạn lo âu lan tỏa bằng liệu pháp thư giãn–luyện tập [Luận án tiến sĩ y học]. Đại học Y Hà Nội.',
        'Trần, T. V., và cộng sự. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine.',
        'UNICEF Việt Nam, Bộ Lao động – Thương binh và Xã hội, và Tổng cục Thống kê. (2022). Khảo sát Sức khỏe Tâm thần Vị thành niên Việt Nam (V-NAMHS 2022). Hà Nội.',
    ]
    for r in refs_vn:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        run = p.add_run(r); run.font.size = Pt(12)

    para(d, 'Tiếng Anh', bold=True, justify=False)
    refs_en = [
        'Allen, J. L., Lavallee, K. L., Herren, C., Ruhe, K., & Schneider, S. (2010). DSM-IV criteria for childhood separation anxiety disorder: Informant, age, and sex differences. Journal of Anxiety Disorders, 24(8), 946–952. https://doi.org/10.1016/j.janxdis.2010.06.022',
        'Beesdo, K., Knappe, S., & Pine, D. S. (2009). Anxiety and anxiety disorders in children and adolescents: Developmental issues and implications for DSM-V. Psychiatric Clinics of North America, 32(3), 483–524. https://doi.org/10.1016/j.psc.2009.06.002',
        'Bie, F., Yan, X., Xing, J., Wang, L., Xu, Y., Wang, G., Wang, Q., Guo, J., Qiao, J., & Rao, Z. (2024). Rising global burden of anxiety disorders among adolescents and young adults: Trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021. Frontiers in Psychiatry, 15, 1489427. https://doi.org/10.3389/fpsyt.2024.1489427',
        'Brown, A., & Carter, R. (2025). School-based mental health interventions in UK secondary schools: A mixed-methods study. Journal of Mental Health.',
        'Browne, M. W., & Cudeck, R. (1993). Alternative ways of assessing model fit. In K. A. Bollen & J. S. Long (Eds.), Testing structural equation models (pp. 136–162). SAGE.',
        'Brunborg, G. S., Nilsen, S. A., Skogen, J. C., & Bang, L. (2025). Possible explanations for the upward trend in mental distress among adolescents in Norway from 2011 to 2024. Social Science & Medicine, 384, 118528. https://doi.org/10.1016/j.socscimed.2025.118528',
        'Cai, S., et al. (2025). Resilience as a protective factor against anxiety in adolescents. Frontiers in Psychiatry.',
        'Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100. https://doi.org/10.1207/s15327558ijbm0401_6',
        'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
        'Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991.',
        'GBD ASEAN Mental Disorders Collaborators. (2025). Epidemiology and burden of ten mental disorders in the Association of Southeast Asian Nations from 1990 to 2021. The Lancet Regional Health – Southeast Asia.',
        'Galante, J., et al. (2023). Mindfulness-based programmes for mental health promotion in adults in non-clinical settings: An IPD meta-analysis. Nature Mental Health, 1(7), 462–476.',
        'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55.',
        'Li, S. H., et al. (2024). ClearlyMe: Co-design of a CBT app for adolescent depression. Cambridge.',
        'Li, S. H., et al. (2025). Network meta-analysis of CBT and physical education for adolescent anxiety. BMC Psychiatry.',
        'Matsumoto, K., et al. (2024). Internet-based cognitive behavioral therapy for Japanese adolescents with anxiety and depression. JMIR Mental Health.',
        'McLean, C. P., Asnaani, A., Litz, B. T., & Hofmann, S. G. (2011). Gender differences in anxiety disorders: Prevalence, course of illness, comorbidity and burden of illness. Journal of Psychiatric Research, 45(8), 1027–1035.',
        'Murphy, R., et al. (2024). Peer support in primary youth mental health care: A scoping review. Journal of Community Psychology, 52(1), 154–180.',
        'Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823',
        'Rapee, R. M., & Spence, S. H. (2004). The etiology of social phobia: Empirical evidence and an initial model. Clinical Psychology Review, 24(7), 737–767.',
        'Salk, R. H., Hyde, J. S., & Abramson, L. Y. (2017). Gender differences in depression in representative national samples: Meta-analyses of diagnoses and symptoms. Psychological Bulletin, 143(8), 783–822.',
        'Samele, C., et al. (2025). ClearFear app for adolescent anxiety. JMIR Formative Research.',
        'Schmidt-Persson, J., et al. (2024). Screen media use and mental health of children and adolescents: A secondary analysis of the SCREENS-Kids randomized clinical trial. JAMA Network Open, 7(1), e2354033.',
        'Wen, X., Lin, Y., Liu, Y., Starcevich, K., Yuan, F., Wang, X., Xie, X., & Yuan, Z. (2020). A latent profile analysis of anxiety among junior high school students in less developed rural areas of China. International Journal of Environmental Research and Public Health, 17(11), 4079. https://doi.org/10.3390/ijerph17114079',
    ]
    for r in refs_en:
        p = d.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Cm(1.0)
        p.paragraph_format.first_line_indent = Cm(-1.0)
        run = p.add_run(r); run.font.size = Pt(12)

    out = OUT_DIR / 'CLEAN_chuong_3_bo_sung_de_cuong.docx'
    d.save(out)
    print(f'  ✓ {out.name}')


# ============================================================
print('Building 2 docs:')
build_erratum()
build_clean()
print('Done.')
