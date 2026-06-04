# -*- coding: utf-8 -*-
"""Kiem tra LA v3_1 - tim cac doan can sua.
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_26052026.docx')

d = Document(LA)
print(f"Total paragraphs: {len(d.paragraphs)}")
print(f"Total tables: {len(d.tables)}")

# Search for key terms
keywords = [
    'bán đô thị',
    'Lê Minh T',
    'Long Bình',
    'Nguyễn Đăng K',
    'Nguyễn Danh Lâm',
    'Bảo Quyên',
    'tổng quát',
    'lan tỏa',
    'Alharbi',
    'Nakie',
    'Jefferies',
]

for kw in keywords:
    count = 0
    positions = []
    for i, p in enumerate(d.paragraphs):
        if kw in p.text:
            count += 1
            if len(positions) < 5:
                positions.append((i, p.text[:200]))
    print(f"\n=== '{kw}': {count} occurrences ===")
    for i, txt in positions:
        print(f"  para {i}: {txt}")
