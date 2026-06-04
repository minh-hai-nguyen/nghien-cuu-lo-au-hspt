# -*- coding: utf-8 -*-
"""Sua loi NBSP (\\xa0) va cross-run spaces.
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

d = Document(FILE)
stats = {'nbsp_punct': 0, 'leading_space_run': 0, 'trailing_space_run': 0, 'cross_double_space': 0,
         'specific_fix': 0}

# =========================================================
# 1. NBSP truoc dau cau ('\xa0:' / '\xa0;' / '\xa0,')
# =========================================================
print("[1] Sua NBSP (\\xa0) truoc dau cau...")
for i, p in enumerate(d.paragraphs):
    for run in p.runs:
        new = run.text
        new = re.sub(r'\xa0([:;,!?])', r'\1', new)  # Remove NBSP before punct
        new = re.sub(r'([:;,!?])\xa0', r'\1 ', new)  # Replace NBSP after with regular space
        if new != run.text:
            run.text = new
            stats['nbsp_punct'] += 1

# =========================================================
# 2. Cross-run space-before-punct: run ends with " ", next run starts with ","
# =========================================================
print("\n[2] Cross-run ' , ' patterns...")
for i, p in enumerate(d.paragraphs):
    runs = p.runs
    for ri in range(len(runs) - 1):
        curr = runs[ri]
        nxt = runs[ri + 1]
        # End of curr is space, start of nxt is punctuation
        if curr.text.endswith(' ') and nxt.text and nxt.text[0] in ',;:!?':
            # Remove trailing space from curr
            curr.text = curr.text.rstrip(' ')
            stats['trailing_space_run'] += 1
        # Start of curr is punctuation, prev ends with space (caught above)

# =========================================================
# 3. Specific cross-run issue: " , Rối loạn..." at start of run
# =========================================================
print("\n[3] Run starts with ' ,' or ' ;' etc...")
for i, p in enumerate(d.paragraphs):
    for run in p.runs:
        # Replace " , Word" at start with ", Word"
        new = re.sub(r'^ ([,;:!?])', r'\1', run.text)
        if new != run.text:
            run.text = new
            stats['leading_space_run'] += 1

# =========================================================
# 4. Cross-run double spaces
# Detection: full text has '  ' but no single run has '  '
# Fix: merge consecutive runs / trim
# =========================================================
print("\n[4] Cross-run double spaces...")
for i, p in enumerate(d.paragraphs):
    if '  ' not in p.text: continue
    runs = p.runs
    for ri in range(len(runs) - 1):
        curr = runs[ri]
        nxt = runs[ri + 1]
        # Both run-internal double-space AND cross-run pattern
        # Pattern: curr ends with ' ', nxt starts with ' ' or curr ends with ' ' + nxt empty/space
        # First: collapse double-internal
        if '  ' in curr.text:
            curr.text = re.sub(r'  +', ' ', curr.text)
            stats['cross_double_space'] += 1
        # Cross: trailing space in curr + leading space in nxt
        while curr.text.endswith(' ') and nxt.text.startswith(' '):
            nxt.text = nxt.text.lstrip(' ')
            if not nxt.text:
                break
            stats['cross_double_space'] += 1
    # Last run
    last = runs[-1]
    if '  ' in last.text:
        last.text = re.sub(r'  +', ' ', last.text)
        stats['cross_double_space'] += 1

# =========================================================
# 5. Trailing spaces at end of paragraph (cosmetic)
# =========================================================
print("\n[5] Trailing/leading spaces in single-space runs at boundaries...")
for i, p in enumerate(d.paragraphs):
    if not p.runs: continue
    # Empty single-space runs at end
    while p.runs and p.runs[-1].text == ' ':
        # If the previous run also ends with non-space, this is redundant trailing
        # Just leave it; it's a single space.
        break

d.save(FILE)

print(f"\nDone. Stats:")
for k, v in stats.items():
    print(f"  {k}: {v}")

# Re-verify
print(f"\n=== VERIFY ===")
d2 = Document(FILE)
remaining_dbl = sum(1 for p in d2.paragraphs if '  ' in p.text)
remaining_nbsp = sum(1 for p in d2.paragraphs if '\xa0' in p.text)
remaining_space_punct = 0
import re
for p in d2.paragraphs:
    if re.search(r'\w\s+[,;:!?]', p.text):
        remaining_space_punct += 1
print(f"  Doan con '  ' (double space): {remaining_dbl}")
print(f"  Doan con '\\xa0' (NBSP): {remaining_nbsp}")
print(f"  Doan con space-before-punct: {remaining_space_punct}")
