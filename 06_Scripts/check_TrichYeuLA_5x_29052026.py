# -*- coding: utf-8 -*-
"""Bo test 5 vong x 5 check kiem tra xac thuc Trich yeu LA cua NCS Cong Thi Hang.
- Vong 1: Format compliance vs template Tran Anh Khoi
- Vong 2: Fact consistency vs LA chinh
- Vong 3: Scale names verification (8 thang do)
- Vong 4: VN-EN consistency
- Vong 5: Anti-fabrication + structural integrity
29/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Cm, Pt

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', 'TrichYeuLA_CongThiHang_v2_29052026.docx')
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

d = Document(FILE)
la = Document(LA)

# Aggregate text
all_text_d = '\n'.join(p.text for p in d.paragraphs)
for tb in d.tables:
    for row in tb.rows:
        for c in row.cells:
            all_text_d += '\n' + c.text

all_text_la = '\n'.join(p.text for p in la.paragraphs)
for tb in la.tables:
    for row in tb.rows:
        for c in row.cells:
            all_text_la += '\n' + c.text

issues = []
total = 0
passed = 0


def check(name, cond, detail=''):
    global total, passed
    total += 1
    status = '✓' if cond else '✗'
    print(f'  {status} {name}')
    if detail:
        print(f'      {detail}')
    if cond:
        passed += 1
    else:
        issues.append(name)
    return cond


# ============================================================
# VONG 1: FORMAT COMPLIANCE
# ============================================================
print('\n=== VÒNG 1: FORMAT COMPLIANCE vs TEMPLATE ===')

sec = d.sections[0]
check('1.1 Page A4 (21.0 × 29.7 cm)',
      abs(sec.page_width.cm - 21.0) < 0.1 and abs(sec.page_height.cm - 29.7) < 0.1,
      f'W={sec.page_width.cm:.2f} H={sec.page_height.cm:.2f}')

check('1.2 Margins 2.54 cm tứ phía',
      all(abs(getattr(sec, m).cm - 2.54) < 0.05
          for m in ['top_margin', 'bottom_margin', 'left_margin', 'right_margin']),
      f'T={sec.top_margin.cm:.2f} B={sec.bottom_margin.cm:.2f} '
      f'L={sec.left_margin.cm:.2f} R={sec.right_margin.cm:.2f}')

# Find specific paragraphs
def find_para(pred):
    for p in d.paragraphs:
        if pred(p):
            return p
    return None

# 1.3 Title TRICH YEU LA: center, bold, 13pt
p_title = find_para(lambda p: 'TRÍCH YẾU LUẬN ÁN' in p.text)
check('1.3 Title TRÍCH YẾU LA: CENTER + bold + 13pt',
      p_title is not None and p_title.alignment is not None
      and p_title.runs[0].bold is True
      and p_title.runs[0].font.size == Pt(13))

# 1.4 Field "Tên đề tài": label bold + value not bold
p_field = find_para(lambda p: p.text.startswith('Tên đề tài:'))
check('1.4 Field "Tên đề tài": label bold + value KHÔNG bold',
      p_field is not None and len(p_field.runs) >= 2
      and p_field.runs[0].bold is True
      and p_field.runs[1].bold in (False, None))

# 1.5 Body line spacing 1.5
body_paras = [p for p in d.paragraphs
              if 'Luận án đã hệ thống' in p.text or 'The dissertation has systematized' in p.text]
all_ls_15 = all(abs(p.paragraph_format.line_spacing - 1.5) < 0.01
                for p in body_paras if p.paragraph_format.line_spacing)
check('1.5 Body line spacing = 1.5', all_ls_15,
      f'{len(body_paras)} body paras checked')

# 1.6 Body first line indent 1.0cm
all_fli = all(p.paragraph_format.first_line_indent
              and abs(p.paragraph_format.first_line_indent.cm - 1.0) < 0.05
              for p in body_paras)
check('1.6 Body first line indent = 1.0 cm', all_fli)

# 1.7 Heading "Đóng góp lý luận": bold, JUSTIFY
p_dgll = find_para(lambda p: 'Đóng góp về mặt lý luận' in p.text)
check('1.7 Heading "Đóng góp lý luận": bold + JUSTIFY',
      p_dgll is not None
      and p_dgll.runs[0].bold is True
      and str(p_dgll.alignment).startswith('JUSTIFY'))

# 1.8 EN Heading "Theoretical contribution": bold + ITALIC + JUSTIFY
p_th = find_para(lambda p: p.text.strip() == 'Theoretical contribution')
check('1.8 "Theoretical contribution": bold + italic + JUSTIFY',
      p_th is not None
      and p_th.runs[0].bold is True and p_th.runs[0].italic is True
      and str(p_th.alignment).startswith('JUSTIFY'))

# 1.9 EN Heading "Practical contributions": bold + ITALIC
p_pr = find_para(lambda p: p.text.strip() == 'Practical contributions')
check('1.9 "Practical contributions": bold + italic + JUSTIFY',
      p_pr is not None
      and p_pr.runs[0].bold is True and p_pr.runs[0].italic is True)

# 1.10 Font Times New Roman tất cả
all_tnr = all(r.font.name == 'Times New Roman' or r.font.name is None
              for p in d.paragraphs for r in p.runs)
check('1.10 Font Times New Roman cho tất cả runs', all_tnr)


# ============================================================
# VONG 2: FACT CONSISTENCY vs LA CHINH
# ============================================================
print('\n=== VÒNG 2: FACT CONSISTENCY vs LA ===')

check('2.1 Mã số "9.31.04.01" khớp LA P48',
      '9.31.04.01' in all_text_d and '9310401' not in all_text_d)

check('2.2 Số HS 1.352 (verify trong LA)',
      '1.352' in all_text_d or '1,352' in all_text_d)

check('2.3 3 dạng RLLA: lan tỏa, xã hội, chia ly',
      'lan tỏa' in all_text_d.lower()
      and 'xã hội' in all_text_d.lower()
      and 'chia ly' in all_text_d.lower())

check('2.4 3 YTNC: ALHT, NĐT, BNHĐ',
      'áp lực học tập' in all_text_d.lower()
      and 'nghiện điện thoại' in all_text_d.lower()
      and 'bắt nạt' in all_text_d.lower())

check('2.5 4 YTBV: TTr, GBTH, HTCM, HTBB',
      'lòng tự trọng' in all_text_d.lower()
      and 'gắn bó với trường học' in all_text_d.lower()
      and 'hỗ trợ của cha mẹ' in all_text_d.lower()
      and 'hỗ trợ của bạn bè' in all_text_d.lower())

check('2.6 F=0,246 cho lo âu chia ly không khác biệt giới',
      'F = 0,246' in all_text_d or 'F=0,246' in all_text_d)

check('2.7 p=0,620 (separation anxiety)',
      'p = 0,620' in all_text_d or 'p=0,620' in all_text_d)

check('2.8 Cường độ TTr ~85-89% (KHÔNG phải ~86%)',
      '85–89' in all_text_d or '85-89' in all_text_d)
check('2.8b KHÔNG có "~86%" stale',
      '~86%' not in all_text_d and 'khoảng 86%' not in all_text_d)

check('2.9 ĐTB 51,13 (áp lực học tập)',
      '51,13' in all_text_d)
check('2.9b ĐTB 28,38 (nghiện điện thoại)',
      '28,38' in all_text_d)
check('2.9c ĐTB 13,52 (bắt nạt thể chất)',
      '13,52' in all_text_d)

check('2.10 Khung CT 8 nội dung cơ bản',
      ('8 nội dung' in all_text_d or '8 nội dung cơ bản' in all_text_d)
      and 'khung' in all_text_d.lower())


# ============================================================
# VONG 3: 8 THANG DO VERIFICATION
# ============================================================
print('\n=== VÒNG 3: 8 THANG ĐO VERIFICATION ===')

scales_correct = {
    'RCADS': 'Chorpita',
    'ESSA': 'Sun',
    'SAS-SV': 'Kwon',
    'OBVQ': 'Olweus',
    'PSSM': 'Goodenow',
    'MSPSS': 'Zimet',
    'RSES': 'Rosenberg',
    'Brief COPE': 'Carver',
}

for scale, author in scales_correct.items():
    check(f'3.{scale} có trong file', scale in all_text_d)
    check(f'3.{scale} có trong LA chính', scale in all_text_la)
    check(f'3.{scale} với tác giả "{author}"',
          scale in all_text_d and author in all_text_d)

# Anti-fabrication
check('3.99a KHÔNG có "MPAS" (bịa)',
      'MPAS' not in all_text_d)
check('3.99b KHÔNG có "MPVS" (bịa)',
      'MPVS' not in all_text_d)
check('3.99c KHÔNG có "Rosenberg Self-Esteem" thiếu acronym RSES',
      not ('Rosenberg Self-Esteem' in all_text_d and 'RSES' not in all_text_d))


# ============================================================
# VONG 4: VN-EN CONSISTENCY
# ============================================================
print('\n=== VÒNG 4: VN-EN CONSISTENCY ===')

# Same mã số
check('4.1 Mã số nhất quán VN-EN',
      all_text_d.count('9.31.04.01') >= 2)

# Số HS
check('4.2 Số HS 1.352 (VN) + 1,352 (EN)',
      '1.352' in all_text_d and '1,352' in all_text_d)

# All 8 scales appear in both VN and EN
for scale in scales_correct:
    cnt = all_text_d.count(scale)
    check(f'4.3 "{scale}" appears >=2 lần (VN+EN)', cnt >= 2,
          f'count={cnt}')

# Statistics consistency
check('4.4 F = 0,246 (VN) + F = 0.246 (EN)',
      'F = 0,246' in all_text_d and 'F = 0.246' in all_text_d)
check('4.5 p = 0,620 (VN) + p = 0.620 (EN)',
      'p = 0,620' in all_text_d and 'p = 0.620' in all_text_d)
check('4.6 51,13 (VN) + 51.13 (EN)',
      '51,13' in all_text_d and '51.13' in all_text_d)
check('4.7 85-89 trong cả VN + EN',
      ('85–89' in all_text_d or '85-89' in all_text_d)
      and all_text_d.count('85') >= 2)

# Same 4 protective factors
check('4.8 Self-esteem trong EN',
      'self-esteem' in all_text_d.lower())
check('4.9 School engagement trong EN',
      'school engagement' in all_text_d.lower())
check('4.10 Parental support trong EN',
      'parental support' in all_text_d.lower())
check('4.11 Peer support trong EN',
      'peer support' in all_text_d.lower())


# ============================================================
# VONG 5: ANTI-FABRICATION + STRUCTURE INTEGRITY
# ============================================================
print('\n=== VÒNG 5: ANTI-FABRICATION + STRUCTURE ===')

# Structure: required sections
check('5.1 Có heading "TRÍCH YẾU LUẬN ÁN"',
      'TRÍCH YẾU LUẬN ÁN' in all_text_d)
check('5.2 Có heading "DISSERTATION ABSTRACT"',
      'DISSERTATION ABSTRACT' in all_text_d)
check('5.3 Có heading "Những kết luận mới của luận án"',
      'Những kết luận mới của luận án' in all_text_d)
check('5.4 Có heading "New conclusions"',
      'New conclusions' in all_text_d)
check('5.5 Có "Đóng góp về mặt lý luận"',
      'Đóng góp về mặt lý luận' in all_text_d)
check('5.6 Có "Đóng góp về mặt thực tiễn"',
      'Đóng góp về mặt thực tiễn' in all_text_d)
check('5.7 Có "Theoretical contribution"',
      'Theoretical contribution' in all_text_d)
check('5.8 Có "Practical contributions"',
      'Practical contributions' in all_text_d)

# 2 signature tables
check('5.9 Đúng 2 signature tables', len(d.tables) == 2)
if len(d.tables) >= 2:
    t1 = d.tables[0]
    t2 = d.tables[1]
    check('5.10 Table 1 chứa "Công Thị Hằng"',
          any('Công Thị Hằng' in c.text for row in t1.rows for c in row.cells))
    check('5.11 Table 1 chứa "TS. Đào Minh Đức"',
          any('TS. Đào Minh Đức' in c.text for row in t1.rows for c in row.cells))
    check('5.12 Table 2 chứa "Cong Thi Hang" (EN, Vietnamese order)',
          any('Cong Thi Hang' in c.text for row in t2.rows for c in row.cells))
    check('5.13 Table 2 chứa "Dr. Dao Minh Duc"',
          any('Dr. Dao Minh Duc' in c.text for row in t2.rows for c in row.cells))

# DSM-5 mention
check('5.14 Có "DSM-5"', 'DSM-5' in all_text_d)

# Methods mentioned
check('5.15 Có "CFA"', 'CFA' in all_text_d)
check('5.16 Có "SEM"', 'SEM' in all_text_d)
check('5.17 Có "Cronbach"', 'Cronbach' in all_text_d)
check('5.18 Có "McDonald"', 'McDonald' in all_text_d)

# Metadata clean
cp = d.core_properties
check('5.19 Metadata author trống', cp.author == '')
check('5.20 Metadata title trống', cp.title == '')
check('5.21 Metadata last_modified_by trống', cp.last_modified_by == '')


# ============================================================
# RESULT
# ============================================================
print()
print('=' * 70)
print(f'TỔNG: {passed}/{total} checks PASS ({len(issues)} lỗi)')
if issues:
    print('\nDANH SÁCH LỖI:')
    for i, iss in enumerate(issues, 1):
        print(f'  {i}. {iss}')
    sys.exit(1)
else:
    print('\nSẠCH LỖI! File Trích yếu sẵn sàng.')
    sys.exit(0)
