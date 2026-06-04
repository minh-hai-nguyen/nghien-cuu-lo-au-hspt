# -*- coding: utf-8 -*-
"""Fix LA v3_5 -> v3_6: Anderson 2025 logic error in para 327.
LA: "tương quan nghịch giữa áp lực học tập và sức khỏe tâm thần kém"
PDF: "positive correlation between academic pressure and POOR mental health"
Su sua: "nghịch" -> "thuận" (pressure UP + poor MH UP = positive correlation).
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_5_FixChen_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_6_FixAnderson_27052026.docx')
RED = RGBColor(192, 0, 0)

REPLACEMENTS = [
    ('mối tương quan nghịch giữa áp lực học tập và sức khỏe tâm thần kém',
     'mối tương quan thuận giữa áp lực học tập và sức khỏe tâm thần kém'),
]


def apply_repl(p, replacements, applied, idx):
    full = p.text
    if not full: return False
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            i = full.find(old, cursor)
            if i == -1: break
            matches.append((i, i + len(old), old, new))
            cursor = i + len(old)
    if not matches: return False
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    safe = []
    last = -1
    for m in matches:
        if m[0] >= last:
            safe.append(m); last = m[1]
    segments = []
    cur = 0
    for s, e, o, n in safe:
        if cur < s:
            segments.append((full[cur:s], False))
        segments.append((n, True))
        applied.append(f"para {idx}: '{o[:60]}' -> '{n[:60]}'")
        cur = e
    if cur < len(full):
        segments.append((full[cur:], False))
    for r in p.runs:
        r.text = ''
    for text, is_red in segments:
        if not text: continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


print(f"=== Fix LA v3_5 -> v3_6 ===")
d = Document(SRC)
applied = []
for i, p in enumerate(d.paragraphs):
    apply_repl(p, REPLACEMENTS, applied, i)
d.save(OUT)
print(f"Saved: {OUT}")
print(f"Total: {len(applied)} fixes")
for a in applied:
    print(f"  ✓ {a}")
