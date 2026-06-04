# -*- coding: utf-8 -*-
"""Kiem tra ban tom tat - 5 vong x 5 check moi vong.
29/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SUM = Document(os.path.join(ROOT, 'Luận án TS', 'TomTatLA_v2_VERIFIED_29052026.docx'))
LA = Document(os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx'))

sum_text = '\n'.join(p.text for p in SUM.paragraphs)
sum_table = '\n'.join(c.text for tb in SUM.tables for row in tb.rows for c in row.cells)
sum_full = sum_text + '\n' + sum_table

la_text = '\n'.join(p.text for p in LA.paragraphs)
la_table = '\n'.join(c.text for tb in LA.tables for row in tb.rows for c in row.cells)
la_full = la_text + '\n' + la_table

issues = []

def check(name, cond, detail=''):
    status = '✓' if cond else '✗'
    print(f'  {status} {name}')
    if detail: print(f'      {detail}')
    if not cond: issues.append(name)
    return cond

# ============================================================
# VONG 1: CAU TRUC (Structure compliance)
# ============================================================
print('=== VONG 1: CAU TRUC ===')
check('1.1 Co MO DAU', 'MỞ ĐẦU' in sum_text)
check('1.2 Co 3 CHUONG', sum_text.count('CHƯƠNG ') >= 3,
      f"Count: {sum_text.count('CHƯƠNG ')}")
check('1.3 Co KET LUAN VA KIEN NGHI', 'KẾT LUẬN VÀ KIẾN NGHỊ' in sum_text)
check('1.4 Co DANH MUC CONG TRINH', 'DANH MỤC CÔNG TRÌNH' in sum_text)
check('1.5 Co Bia + Phu Bia', 'BỘ GIÁO DỤC' in sum_text and 'Phản biện 1' in sum_text)

# ============================================================
# VONG 2: NOI DUNG MO DAU (10 MUC)
# ============================================================
print('\n=== VONG 2: 10 MUC MO DAU ===')
mucs = ['1. Lý do chọn đề tài', '2. Mục đích nghiên cứu', '3. Đối tượng nghiên cứu',
        '4. Khách thể nghiên cứu', '5. Giả thuyết khoa học', '6. Nhiệm vụ nghiên cứu',
        '7. Giới hạn nghiên cứu', '8. Phương pháp nghiên cứu',
        '9. Đóng góp mới', '10. Cấu trúc luận án']
for m in mucs:
    check(f'2.{mucs.index(m)+1} {m}', m in sum_text)

# ============================================================
# VONG 3: SO LIEU CHU CHOT MATCH LA
# ============================================================
print('\n=== VONG 3: SO LIEU CHU CHOT MATCH LA ===')
key_facts = [
    ('1.352', '1.352 mau'),
    ('614', 'Nam 614'),
    ('738', 'Nu 738'),
    ('51,13', 'ALHT DTB'),
    ('28,38', 'NDT DTB'),
    ('13,52', 'BNHD DTB'),
]
for fact, desc in key_facts:
    in_sum = fact in sum_full
    in_la = fact in la_full
    check(f'3.X {desc}: {fact}', in_sum and in_la,
          f'In SUM: {in_sum}, In LA: {in_la}')

# ============================================================
# VONG 4: BETA + ALPHA VALUES MATCH LA
# ============================================================
print('\n=== VONG 4: BETA + ALPHA VALUES ===')
key_values = [
    ('0,510', 'ALHT->RLLALT beta'),
    ('0,533', 'ALHT->RLLA(3f)'),
    ('0,400', 'NDT->RLLA'),
    ('0,276', 'BNHD->RLLA'),
    ('-0,457', 'TTr->RLLA'),
    ('0,811', 'RLLALT alpha'),
    ('0,708', 'ALHT alpha'),
    ('0,836', 'NDT alpha'),
    ('0,725', 'TTr alpha'),
]
for v, desc in key_values:
    in_sum = v in sum_full
    in_la = v in la_full
    check(f'4.X {desc}: {v}', in_sum and in_la,
          f'In SUM: {in_sum}, In LA: {in_la}')

# ============================================================
# VONG 5: CITATION + VAN PHONG
# ============================================================
print('\n=== VONG 5: CITATION + VAN PHONG ===')
check('5.1 V-NAMHS 5.996', '5.996' in sum_full)
check('5.2 V-NAMHS 2,3%', '2,3%' in sum_full)
check('5.3 Xu 373.216', '373.216' in sum_full)
check('5.4 Compas N=80.850', '80.850' in sum_full)
check('5.5 Steare 48/52', '48 nghiên cứu' in sum_full and '52 nghiên cứu' in sum_full)
check('5.6 Saikia 2023', 'Saikia' in sum_full)
check('5.7 Chen 2023', 'Chen' in sum_full)
check('5.8 Hoàng Trung Học 2025', 'Hoàng Trung Học' in sum_full)
check('5.9 Bảo Quyên 2025', 'Bảo Quyên' in sum_full)
check('5.10 Không dùng "có lẽ"', 'có lẽ' not in sum_text)
check('5.11 Không dùng "có thể là"', 'có thể là' not in sum_text)

# ============================================================
# Tổng kết
# ============================================================
print()
print('=' * 70)
print(f'TONG: {len(issues)} loi')
if issues:
    print('\nDanh sach loi:')
    for i, iss in enumerate(issues):
        print(f'  {i+1}. {iss}')
else:
    print('SACH LOI')
