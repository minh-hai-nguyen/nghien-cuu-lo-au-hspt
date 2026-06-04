# -*- coding: utf-8 -*-
"""Kiem tra LAN CUOI tung dong - tat ca cac loai loi co the.
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

print("="*80)
print("KIEM TRA CUOI - 01_LuanAn_v3_1_FixCoping_28052026.docx")
print("="*80)

results = []

# 1. Double spaces in body text
dbl = sum(1 for p in d.paragraphs if '  ' in p.text)
results.append(('Khoang trang doi (  )', dbl, 0))

# 2. Space before punctuation (excluding intentional like "p < 0.001")
space_before = 0
for p in d.paragraphs:
    if re.search(r'\w\s+[,;:!?](?!\d)', p.text):
        space_before += 1
results.append(('Dau cach truoc dau cau (sai)', space_before, 0))

# 3. NBSP next to punctuation
nbsp_punct = sum(1 for p in d.paragraphs if re.search(r'\xa0[,;:!?]|[,;:!?]\xa0', p.text))
results.append(('NBSP canh dau cau', nbsp_punct, 0))

# 4. Heading numbering duplicates (within same section)
heading_nums = {}
for i, p in enumerate(d.paragraphs):
    if 'Heading' not in p.style.name and 'Tiêu đề' not in p.style.name:
        continue
    text = p.text.strip()
    m = re.match(r'^(\d+(?:\.\d+)*)', text)
    if m:
        num = m.group(1)
        heading_nums.setdefault(num, []).append((i, text[:60]))

# Real dupes = same number in same parent + same chapter context
# Find true duplicates (same number, in same section context)
true_dupes = []
for num, lst in heading_nums.items():
    if len(lst) <= 1: continue
    # Check if these are actually different sections (CHƯƠNG vs KẾT LUẬN)
    # If para distance > 200, they might be in different chapters/sections
    sorted_lst = sorted(lst)
    for j in range(1, len(sorted_lst)):
        if sorted_lst[j][0] - sorted_lst[j-1][0] < 200:
            true_dupes.append((num, sorted_lst[j-1], sorted_lst[j]))
results.append(('Tieu de SO TRUNG trong cung section', len(true_dupes), 0))
if true_dupes:
    print("\nTrue duplicates:")
    for num, a, b in true_dupes:
        print(f"  {num}: para {a[0]} '{a[1]}' vs para {b[0]} '{b[1]}'")

# 5. Diacritic inconsistency
hoa_count = sum(p.text.count('hoà') for p in d.paragraphs)
hoa2_count = sum(p.text.count('hòa') for p in d.paragraphs)
results.append(("Chinh ta 'hoà' (chua sua)", hoa_count, 0))

tt1 = sum(p.text.count('Tâm Thần') for p in d.paragraphs)
results.append(("'Tâm Thần' (sai capitalize)", tt1, 0))

# 6. Numbering inconsistencies (1.X duplicate IN SAME CHAPTER)
# Check by counting heading sequences
groups = {}
for i, p in enumerate(d.paragraphs):
    if 'Heading' not in p.style.name and 'Tiêu đề' not in p.style.name:
        continue
    text = p.text.strip()
    m = re.match(r'^(\d+(?:\.\d+)*)', text)
    if m:
        num = m.group(1)
        parts = num.split('.')
        if len(parts) >= 2:
            parent = '.'.join(parts[:-1])
            try:
                groups.setdefault(parent, []).append((int(parts[-1]), i, num, text[:60]))
            except: pass

seq_issues = 0
for parent, children in groups.items():
    children.sort()
    nums = [c[0] for c in children]
    if any(nums[j] - nums[j-1] not in (0, 1) for j in range(1, len(nums))):
        seq_issues += 1  # Gap or duplicate
    # Check for consecutive duplicates (within same section)
    for j in range(1, len(children)):
        if children[j][0] == children[j-1][0] and children[j][1] - children[j-1][1] < 100:
            seq_issues += 1
results.append(('So thu tu heading nhay/trung', seq_issues, 0))

# 7. Text integrity - count chars
total_chars = sum(len(p.text) for p in d.paragraphs)
total_paras = len(d.paragraphs)
total_tables = len(d.tables)
print(f"\n--- Thong ke ---")
print(f"  Paragraphs: {total_paras}")
print(f"  Tables: {total_tables}")
print(f"  Total chars: {total_chars}")
print(f"  File size: {os.path.getsize(FILE)//1024} KB")

# 8. Verify priority fixes still in place
print(f"\n--- Verify 5 lỗi ưu tiên (đã sửa) ---")
priority_checks = [
    (149, 'và cs.', 'cs..', '[1.149] cs.. -> cs.'),
    (331, '1.1.6.5', '1.2.6.5', '[2.331] 1.2.6.5 -> 1.1.6.5'),
    (544, '1.2.4.4', None, '[3.544] 1.2.4.3 -> 1.2.4.4 (cascade)'),
    (549, '1.2.4.5', None, '[3b.549] 1.2.4.4 -> 1.2.4.5 (cascade)'),
    (1002, '3.4.2.4', None, '[4.1002] 3.4.2.1 -> 3.4.2.4'),
    (1036, '3.4.3.3', None, '[5.1036] 3.4.3.2 -> 3.4.3.3'),
]
for idx, want, not_want, desc in priority_checks:
    if idx >= len(d.paragraphs): continue
    text = d.paragraphs[idx].text
    ok = want in text and (not_want is None or text.find(not_want) > text.find(want))
    mark = '✓' if ok else '✗'
    print(f"  {mark} {desc}")

# Final summary
print()
print("="*80)
print("KET QUA")
print("="*80)
for label, n, expected in results:
    mark = '✓' if n == expected else '✗'
    print(f"  {mark} {label}: {n} (expected {expected})")

# Bracket balance still
unbalanced = sum(1 for p in d.paragraphs if p.text.count('(') != p.text.count(')'))
print(f"\n  CHU Y (chua sua, can NCS xem): {unbalanced} doan co dau ngoac khong can bang")
print(f"  CHU Y (chua sua): 'hoà' van con {hoa_count} - kiem tra lai")
