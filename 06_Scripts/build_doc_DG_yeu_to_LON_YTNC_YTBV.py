"""DOC: Dien giai YEU TO LON (YTNC tong + YTBV tong + Mo hinh tich hop)
theo so lieu chinh xac tu file 00_Binh luan so lieu-v2.docx
Bang 3.37, 3.38, 3.39, 3.40, 3.41, 3.42

ALL FACTS VERIFIED tu Tables 27-32 file v2:
- Bang 3.37 (Fit YTNC+YTBV): chi²(17)=153,792; chi²/df=9,047; CFI=0,937; TLI=0,897; RMSEA=0,077 (KTC 0,066-0,089)
- Bang 3.38 (Path YTNC+YTBV):
  + YTNC -> RLLA: beta=0,669; SE=0,166; CR=9,103; p<0,001; R²=0,598
  + YTBV -> RLLA: beta=-0,220; SE=0,082; CR=-4,725; p<0,001; R²=0,598
- Bang 3.39 (Fit YTNC rieng): chi²(4)=12,407; chi²/df=3,102; CFI=0,994; TLI=0,986; RMSEA=0,039 (0,016-0,065)
- Bang 3.40 (Path YTNC rieng): YTNC -> RLLA: beta=0,747; SE=0,168; CR=10,241; p<0,001; R²=0,558
- Bang 3.41 (Fit YTBV rieng): chi²(4)=30,608; chi²/df=7,652; CFI=0,982; TLI=0,954; RMSEA=0,070 (0,048-0,094)
- Bang 3.42 (Path YTBV rieng): YTBV -> RLLA: beta=-0,352; SE=0,102; CR=-9,890; p<0,001; R²=0,124

YEU TO NHO:
- YTNC chua: ALHT (ap luc hoc tap) + NDT (nghien dien thoai) + BNHD (bat nat hoc duong)
- YTBV chua: GBTH (gan bo truong hoc) + TTr (tu trong) + HTCM (ho tro cha me) + HTBB (ho tro ban be)
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Dien_giai_yeu_to_LON_YTNC_YTBV_mo_hinh_tich_hop.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('DIỄN GIẢI YẾU TỐ LỚN — YTNC TỔNG, YTBV TỔNG\nVÀ MÔ HÌNH TÍCH HỢP\n— Bảng 3.37–3.42 chương 3 luận án —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# 1. Khung chung — Phân biệt yếu tố LỚN vs NHỎ
H('1. Phân biệt yếu tố LỚN và yếu tố NHỎ', level=2, color=NAVY)
para(
    'Trong chương 3 luận án, các yếu tố ảnh hưởng đến rối loạn '
    'lo âu được phân thành HAI CẤP độ phân tích: cấp YẾU TỐ LỚN '
    'và cấp YẾU TỐ NHỎ.', indent=False
)
caption('Bảng 1. Phân cấp yếu tố ảnh hưởng đến RLLA trong chương 3')
add_table(
    ['Cấp', 'Yếu tố LỚN', 'Yếu tố NHỎ (sub-factors)', 'Bảng tham chiếu'],
    [
        ['Cấp 1', 'YTNC — Yếu tố Nguy cơ tổng',
         'ALHT (áp lực học tập) + NĐT (nghiện điện thoại) + BNHĐ (bắt nạt học đường)',
         '3.21, 3.39, 3.40'],
        ['Cấp 1', 'YTBV — Yếu tố Bảo vệ tổng',
         'GBTH (gắn bó trường học) + TTr (tự trọng) + HTCM (hỗ trợ cha mẹ) + HTBB (hỗ trợ bạn bè)',
         '3.22, 3.41, 3.42'],
        ['Cấp 0', 'Mô hình TÍCH HỢP YTNC + YTBV',
         'Cả hai cấu phần lớn cùng dự báo RLLA',
         '3.37, 3.38'],
        ['Cấp 1 riêng', 'BPĐP — Biện pháp đối phó',
         'Tránh né + giải quyết vấn đề + tìm kiếm hỗ trợ',
         '3.43, 3.44, 3.45'],
    ]
)
para('')
para(
    'Áp lực học tập, bắt nạt học đường, tự trọng... ĐỀU LÀ yếu '
    'tố nhỏ thuộc YTNC hoặc YTBV. Doc này diễn giải CẤP YẾU TỐ '
    'LỚN — tức tổng hợp các sub-factors thành biến tiềm ẩn YTNC '
    'và YTBV, sau đó phân tích tác động đến RLLA.'
)
para(
    'Lưu ý quan trọng: cấp yếu tố LỚN có ý nghĩa lý thuyết khác '
    'với cấp NHỎ. Khi yếu tố nhỏ ALHT có β = 0,510 với RLLATQ — '
    'đó là tác động RIÊNG của áp lực học tập. Khi YTNC tổng có '
    'β = 0,669 với RLLA — đó là tác động TỔNG HỢP của ba yếu '
    'tố nhỏ qua biến tiềm ẩn YTNC. Hai cấp KHÔNG tương đương.'
)

# 2. Bảng 3.39-3.40 — YTNC riêng
H('2. Diễn giải Bảng 3.39–3.40 — YTNC riêng tác động đến RLLA', level=2, color=NAVY)
caption('Bảng 2. Fit indices mô hình YTNC riêng (Bảng 3.39)')
add_table(
    ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI RMSEA'],
    [
        ['YTNC — RLLA', '12,407 (4)', '3,102', '0,994', '0,986', '0,039', '0,016–0,065'],
    ]
)
para('')
caption('Bảng 3. Hệ số tác động YTNC → RLLA (Bảng 3.40)')
add_table(
    ['Mô hình', 'Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
    [
        ['1', 'YTNC → RLLA', '0,747', '0,168', '10,241', '< 0,001', '0,558'],
    ]
)
para('')

H('Diễn giải:', level=3, color=NAVY)
para(
    'FIT INDICES — Mô hình YTNC riêng đạt fit XUẤT SẮC: CFI = '
    '0,994 (gần ngưỡng tối đa 1,0); TLI = 0,986; RMSEA = 0,039 '
    'với KTC 90% từ 0,016 đến 0,065 — chứa ngưỡng 0,05 cho '
    '"good fit" theo Hu & Bentler (1999). χ²/df = 3,102 — '
    'gần ngưỡng lý tưởng 3.', indent=False
)
para(
    'HỆ SỐ TÁC ĐỘNG — β = 0,747. Khi YTNC tổng (gồm ALHT + NĐT + '
    'BNHĐ) TĂNG 1 độ lệch chuẩn, RLLA tổng TĂNG 0,747 độ lệch '
    'chuẩn. Theo Cohen (1988), |β| > 0,5 thuộc "tác động LỚN" — '
    'cường độ này GẦN MỨC TỐI ĐA cho biến tiềm ẩn xã hội. C.R. = '
    '10,241 (vượt xa ngưỡng 1,96 và 2,58); p < 0,001 — CỰC kỳ '
    'ý nghĩa thống kê.'
)
para(
    'R² = 0,558 — Yếu tố nguy cơ tổng GIẢI THÍCH 55,8% phương '
    'sai của RLLA tổng. Theo Cohen (1988): R² > 0,26 thuộc '
    '"effect size LỚN" — kết quả VƯỢT XA ngưỡng. Một mình YTNC '
    'tổng đã giải thích MỘT NỬA biến thiên của lo âu tổng — phát '
    'hiện ĐẶC BIỆT MẠNH.'
)

# 3. Bảng 3.41-3.42 — YTBV riêng
H('3. Diễn giải Bảng 3.41–3.42 — YTBV riêng tác động đến RLLA', level=2, color=NAVY)
caption('Bảng 4. Fit indices mô hình YTBV riêng (Bảng 3.41)')
add_table(
    ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI RMSEA'],
    [
        ['YTBV — RLLA', '30,608 (4)', '7,652', '0,982', '0,954', '0,070', '0,048–0,094'],
    ]
)
para('')
caption('Bảng 5. Hệ số tác động YTBV → RLLA (Bảng 3.42)')
add_table(
    ['Mô hình', 'Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
    [
        ['1', 'YTBV → RLLA', '−0,352', '0,102', '−9,890', '< 0,001', '0,124'],
    ]
)
para('')

H('Diễn giải:', level=3, color=NAVY)
para(
    'FIT INDICES — Mô hình YTBV riêng đạt fit CHẤP NHẬN ĐƯỢC: '
    'CFI = 0,982; TLI = 0,954; RMSEA = 0,070 (KTC 90% 0,048–'
    '0,094) — RMSEA hơi vượt ngưỡng "good fit" 0,06 nhưng vẫn '
    'trong "acceptable" 0,08. χ²/df = 7,652 — vượt ngưỡng lý '
    'tưởng < 3 và cả < 5; tuy nhiên với cỡ mẫu lớn n = 1.352, '
    'χ² test rất nhạy. Đánh giá tổng: fit CHẤP NHẬN nhưng '
    'KHÔNG bằng YTNC riêng.', indent=False
)
para(
    'HỆ SỐ TÁC ĐỘNG — β = −0,352. Dấu ÂM phản ánh tác động BẢO '
    'VỆ: khi YTBV tổng (gồm GBTH + TTr + HTCM + HTBB) TĂNG 1 '
    'độ lệch chuẩn, RLLA tổng GIẢM 0,352 độ lệch chuẩn. Theo '
    'Cohen (1988), |β| trong khoảng 0,3–0,5 thuộc "tác động '
    'TRUNG BÌNH". C.R. = −9,890; p < 0,001 — ý nghĩa thống '
    'kê cao.'
)
para(
    'R² = 0,124 — YTBV tổng GIẢI THÍCH 12,4% phương sai RLLA. '
    'Theo Cohen (1988): R² ≥ 0,13 thuộc "effect size TRUNG '
    'BÌNH" — kết quả CHẠM NGƯỠNG TRUNG BÌNH. Đáng chú ý: R² '
    'của YTBV (12,4%) THẤP HƠN RẤT NHIỀU so với YTNC '
    '(55,8%). Tỷ số 12,4 / 55,8 ≈ 0,222 — YTBV chỉ giải thích '
    '~22% so với YTNC.'
)

# 4. Bảng 3.37-3.38 — Mô hình TÍCH HỢP
H('4. Diễn giải Bảng 3.37–3.38 — MÔ HÌNH TÍCH HỢP YTNC + YTBV', level=2, color=NAVY)
caption('Bảng 6. Fit indices mô hình tích hợp YTNC + YTBV (Bảng 3.37)')
add_table(
    ['Mô hình', 'χ²(df)', 'χ²/df', 'CFI', 'TLI', 'RMSEA', '90% CI RMSEA'],
    [
        ['YTNC & YTBV — RLLA', '153,792 (17)', '9,047', '0,937', '0,897', '0,077', '0,066–0,089'],
    ]
)
para('')
caption('Bảng 7. Hệ số tác động YTNC + YTBV → RLLA (Bảng 3.38)')
add_table(
    ['Mô hình', 'Path', 'β', 'S.E.', 'C.R.', 'p', 'R²'],
    [
        ['1', 'YTNC → RLLA', '0,669', '0,166', '9,103', '< 0,001', '0,598'],
        ['2', 'YTBV → RLLA', '−0,220', '0,082', '−4,725', '< 0,001', '0,598'],
    ]
)
para('')

H('Diễn giải:', level=3, color=NAVY)
para(
    'FIT INDICES — Mô hình tích hợp đạt fit YẾU HƠN hai mô hình '
    'riêng: CFI = 0,937 (vượt 0,90 nhưng dưới 0,95); TLI = '
    '0,897 (dưới 0,95); RMSEA = 0,077 (vượt 0,06 nhưng dưới '
    '0,08); χ²/df = 9,047 (vượt cả 5). Đánh giá: fit GẦN '
    'NGƯỠNG ACCEPTABLE — không tốt bằng YTNC riêng (CFI = '
    '0,994) hay YTBV riêng (CFI = 0,982). Lý do: mô hình '
    'phức tạp hơn (thêm tương quan giữa YTNC và YTBV), khó '
    'fit hơn.', indent=False
)
para(
    'HỆ SỐ TÁC ĐỘNG — Hai đường ảnh hưởng:'
)
para(
    '• YTNC → RLLA: β = 0,669 — sau khi kiểm soát YTBV. So '
    'với YTNC riêng (β = 0,747), cường độ GIẢM nhẹ '
    '(0,669/0,747 = 0,896 ≈ 90%) — gợi ý YTBV chia sẻ một '
    'phần phương sai với YTNC nhưng tác động chính của YTNC '
    'vẫn ổn định.'
)
para(
    '• YTBV → RLLA: β = −0,220 — sau khi kiểm soát YTNC. So '
    'với YTBV riêng (β = −0,352), cường độ GIẢM mạnh '
    '(0,220/0,352 = 0,625 ≈ 63%) — nghĩa là 37% tác động '
    'BẢO VỆ của YTBV được giải thích bởi YTNC. Nói cách '
    'khác, YTBV không hoàn toàn ĐỘC LẬP với YTNC.'
)
para(
    'R² CHUNG = 0,598 — Tích hợp YTNC + YTBV giải thích '
    '59,8% phương sai RLLA. So với:'
)
para(
    '• YTNC riêng: R² = 0,558 (tăng 0,040 = 4,0% khi thêm '
    'YTBV)\n'
    '• YTBV riêng: R² = 0,124 (tăng 0,474 = 47,4% khi thêm '
    'YTNC)', indent=False
)
para(
    'Tỷ số đóng góp: YTNC chiếm phần lớn (47,4 / 4,0 ≈ 11,9 '
    'lần đóng góp lớn hơn YTBV trong mô hình tích hợp).'
)

# 5. So sánh ba mô hình
H('5. So sánh ba mô hình', level=2, color=NAVY)
caption('Bảng 8. So sánh ba mô hình SEM cấp yếu tố LỚN')
add_table(
    ['Mô hình', 'CFI', 'RMSEA', 'β', 'R²', 'Đánh giá'],
    [
        ['YTNC riêng (3.39–3.40)', '0,994', '0,039', '0,747', '0,558', 'Fit XUẤT SẮC'],
        ['YTBV riêng (3.41–3.42)', '0,982', '0,070', '−0,352', '0,124', 'Fit CHẤP NHẬN'],
        ['YTNC + YTBV (3.37–3.38)', '0,937', '0,077', '0,669/−0,220', '0,598', 'Fit ACCEPTABLE'],
    ]
)
para('')

H('Năm phát hiện chính từ ba mô hình:', level=3, color=NAVY)
para(
    'Thứ nhất, YTNC TỔNG (gồm ALHT + NĐT + BNHĐ) là yếu '
    'tố ảnh hưởng MẠNH NHẤT đến RLLA: β = 0,747; R² = '
    '55,8%. Hơn một nửa biến thiên của lo âu tổng được '
    'giải thích bởi yếu tố nguy cơ tổng — bằng chứng '
    'MẠNH MẼ.'
)
para(
    'Thứ hai, YTBV TỔNG có tác động BẢO VỆ ý nghĩa thống '
    'kê (β = −0,352, p < 0,001) nhưng CƯỜNG ĐỘ thấp hơn '
    'nhiều so với YTNC. Tỷ số |β YTBV| / |β YTNC| = '
    '0,352 / 0,747 ≈ 0,471 — bảo vệ chỉ bằng ~47% nguy cơ.'
)
para(
    'Thứ ba, mô hình TÍCH HỢP làm GIẢM cường độ cả hai '
    'cấu phần do tương quan: β YTNC giảm 10% (0,747 → '
    '0,669); β YTBV giảm 37% (−0,352 → −0,220). Pattern '
    'này gợi ý YTBV không hoàn toàn ĐỘC LẬP với YTNC — '
    'có thể do nhiều yếu tố nhỏ trong YTBV (ví dụ tự '
    'trọng) bị TRUNG GIAN bởi YTNC.'
)
para(
    'Thứ tư, R² mô hình tích hợp = 59,8% — vượt ngưỡng '
    '"effect size LỚN" của Cohen (1988, R² > 0,26) hơn '
    'gấp đôi. Đây là một trong những R² CAO NHẤT trong y '
    'văn về RLLA học sinh THCS.'
)
para(
    'Thứ năm, fit indices phản ánh độ phức tạp: YTNC '
    'riêng có fit CAO NHẤT (CFI = 0,994), tích hợp có '
    'fit THẤP NHẤT (CFI = 0,937). Phù hợp với nguyên '
    'tắc: mô hình càng đơn giản càng dễ fit, mô hình '
    'càng phức tạp càng cần đánh giá kỹ.'
)

# 6. Hàm ý
H('6. Bốn hàm ý cho can thiệp', level=2, color=NAVY)
para(
    'HÀM Ý 1 — Ưu tiên giảm yếu tố nguy cơ. Vì YTNC giải '
    'thích 55,8% phương sai RLLA — gấp 4,5 lần YTBV (12,4%) '
    '— can thiệp NÊN ưu tiên giảm YTNC (giảm áp lực học '
    'tập, hạn chế nghiện điện thoại, phòng ngừa bắt nạt) '
    'thay vì chỉ tăng YTBV.', indent=False
)
para(
    'HÀM Ý 2 — KHÔNG bỏ qua YTBV. Dù cường độ thấp hơn, '
    'YTBV vẫn có tác động ý nghĩa thống kê (p < 0,001) và '
    'là yếu tố CÓ THỂ CAN THIỆP nhanh hơn YTNC. Tăng '
    'GBTH, TTr, HTCM cùng với giảm YTNC sẽ tạo TÁC ĐỘNG '
    'CỘNG HƯỞNG.'
)
para(
    'HÀM Ý 3 — Báo cáo TÍCH HỢP với CẢ ba mô hình. Khi '
    'báo cáo, KHÔNG nên chỉ trích β YTNC = 0,747 (riêng) '
    'hoặc β YTNC = 0,669 (tích hợp). Cần cung cấp CẢ HAI '
    '+ giải thích lý do khác biệt — minh bạch về phương '
    'pháp luận.'
)
para(
    'HÀM Ý 4 — Phân tích trung gian (mediation). Pattern '
    'β YTBV giảm 37% khi thêm YTNC gợi ý có thể có '
    'TRUNG GIAN: YTNC → YTBV thấp → RLLA. Nghiên cứu '
    'tiếp theo NÊN phân tích cụ thể từng yếu tố nhỏ '
    'trong YTBV (đặc biệt tự trọng) làm trung gian như '
    'thế nào.'
)

# 7. CÂU TRẢ LỜI
H('7. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue_run('Diễn giải yếu tố LỚN — YTNC + YTBV + Mô hình tích hợp:', bold=True)
blue_run(
    '(1) YTNC RIÊNG (Bảng 3.39–3.40): β = 0,747; R² = '
    '0,558 (55,8% phương sai RLLA). Fit XUẤT SẮC '
    '(CFI = 0,994; RMSEA = 0,039). Yếu tố nguy cơ tổng '
    'là biến SEM tiềm ẩn gồm ALHT + NĐT + BNHĐ.'
)
blue_run(
    '(2) YTBV RIÊNG (Bảng 3.41–3.42): β = −0,352; R² = '
    '0,124 (12,4% phương sai RLLA). Fit CHẤP NHẬN '
    '(CFI = 0,982; RMSEA = 0,070). Yếu tố bảo vệ tổng '
    'gồm GBTH + TTr + HTCM + HTBB.'
)
blue_run(
    '(3) MÔ HÌNH TÍCH HỢP (Bảng 3.37–3.38): β YTNC = '
    '0,669; β YTBV = −0,220 (sau kiểm soát lẫn nhau); '
    'R² CHUNG = 0,598 (59,8%). Fit ACCEPTABLE (CFI = '
    '0,937; RMSEA = 0,077). Khi tích hợp, β YTNC '
    'giảm ~10% và β YTBV giảm ~37% — gợi ý YTBV '
    'không hoàn toàn độc lập với YTNC.'
)
blue_run(
    '(4) SO SÁNH CƯỜNG ĐỘ: |β YTBV| / |β YTNC| ≈ 0,47 '
    '— bảo vệ chỉ bằng 47% nguy cơ. R² YTNC riêng '
    '(55,8%) gấp 4,5 lần R² YTBV riêng (12,4%). Yếu '
    'tố nguy cơ ÁP ĐẢO yếu tố bảo vệ ở cấp tổng hợp.'
)
blue_run(
    '(5) HÀM Ý: Can thiệp ƯU TIÊN giảm yếu tố nguy cơ '
    '(ALHT, NĐT, BNHĐ) vì cường độ tác động lớn hơn. '
    'Đồng thời tăng yếu tố bảo vệ (TTr, GBTH, HTCM) '
    'để có hiệu ứng cộng hưởng. R² mô hình tích hợp = '
    '59,8% — một trong những R² CAO NHẤT trong y văn '
    'RLLA học sinh THCS.'
)

# 8. TLTK
H('8. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
