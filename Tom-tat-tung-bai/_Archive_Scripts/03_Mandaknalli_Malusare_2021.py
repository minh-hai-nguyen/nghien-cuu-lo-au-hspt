# -*- coding: utf-8 -*-
"""Bài 3: Mandaknalli & Malusare 2021 — theo phong cách CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 3 =====
bl('Tóm tắt bài 3', bold=True)

# BƯỚC 1: ĐỊNH DANH
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình nghiên cứu \u00ab Nghiên cứu cắt ngang về tỷ lệ lo âu ở khu vực trường học thuộc đô thị tự trị \u00bb (A Cross-Sectional Study on the Prevalence of Anxiety among Municipality School Area), do Rahul Mandaknalli và Ragini Malusare (2021), thuộc Khoa Tâm thần học, Trường Cao đẳng Y khoa Mahadevappa Rampure, Ấn Độ, khảo sát trên mẫu 450 học sinh từ lớp 9 đến lớp 12 tại 3 trường thuộc khu vực đô thị tự trị, xuất bản trên tạp chí MedPulse International Journal of Psychology, tập 18, số 3, tháng 6/2021.')

# BƯỚC 2: TỔNG QUAN PP + "Nói cách khác" (KT1)
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng Thang đo Rối loạn Lo âu Tổng quát 7 mục (GAD-7 \u2014 Generalized Anxiety Disorder 7-item Scale) phiên bản ngôn ngữ địa phương kết hợp với bảng câu hỏi tự báo cáo về nhân khẩu học-xã hội và lối sống. Nói cách khác, công trình này sử dụng phương pháp sàng lọc cắt ngang tại trường học bằng bảng câu hỏi.', bold=True)

# BƯỚC 3: ĐỊNH NGHĨA (KT5 — phả hệ)
bl('GAD-7 là thang đo tự báo cáo 7 mục do Spitzer và cộng sự (2006) phát triển, được sử dụng rộng rãi để sàng lọc rối loạn lo âu tổng quát. Thang điểm từ 0 đến 21, phân loại thành bốn mức: tối thiểu (0\u20134), nhẹ (5\u20139), trung bình (10\u201314) và nặng (15\u201321).')

# BƯỚC 4: BIỆN MINH (KT4)
bl('Tổng quan tài liệu của chúng tôi cho thấy tỷ lệ gộp toàn cầu của các rối loạn tâm thần ở trẻ em và thanh thiếu niên là 13,4% (Polanczyk và cộng sự, 2015), và tỷ lệ suốt đời của bất kỳ rối loạn lo âu nào ở trẻ em và thanh thiếu niên là khoảng 20% (Pine và cộng sự, 2009). Tuy nhiên, dữ liệu tại các trường đô thị tự trị ở Ấn Độ còn hạn chế.', bold=True)

# BƯỚC 5: LIỆT KÊ DỮ LIỆU
bl('Dữ liệu nhân khẩu học-xã hội bao gồm tuổi, giới tính, thu nhập hàng tháng của gia đình, loại hình gia đình (hạt nhân hoặc đại gia đình), khối lớp học, tình trạng sống cùng gia đình và cảm nhận về mối quan hệ với bạn bè. Thông tin lối sống bao gồm hoạt động thể chất thường xuyên, mức độ hài lòng với giấc ngủ, số giờ ngủ mỗi đêm, hút thuốc lá và sử dụng internet.')

# BƯỚC 6: MÔ TẢ QUY TRÌNH
bl('Nghiên cứu được thực hiện trong thời gian 1 năm dưới sự quản lý của Khoa Tâm thần học, Trường Cao đẳng Y khoa Mahadevappa Rampure. Ba trường được xác định và 450 học sinh từ lớp 9 đến lớp 12 tham gia nghiên cứu. Một kiểm tra thử nghiệm (pilot test) được tiến hành trước trên 30 học sinh để xác nhận độ tin cậy của bảng câu hỏi. Sự chấp thuận đạo đức đã được lấy từ ủy ban đạo đức cơ sở địa phương. Dữ liệu được phân tích bằng SPSS phiên bản 23.0 với kiểm định chi bình phương hoặc kiểm định chính xác Fisher, giá trị P < 0,05 được coi là có ý nghĩa thống kê.')

# BƯỚC 9: ĐỊNH LƯỢNG — BẢNG
rb('Kết quả nghiên cứu định lượng')

bl('Tỷ lệ lo âu và phân bố theo mức độ nghiêm trọng (N = 450, GAD-7):', bold=True)
tbl(
    ['Mức độ lo âu', 'Số học sinh', 'Tỷ lệ (trong nhóm lo âu)'],
    [
        ['Tổng có lo âu', '108/450', '24,0%'],
        ['Nhẹ (GAD 0\u20134)', '34', '31,48%'],
        ['Trung bình (GAD 5\u20139)', '43', '39,81%'],
        ['Nặng (GAD 10\u201314)', '17', '15,74%'],
        ['Rất nặng (GAD 15\u201321)', '14', '12,96%'],
    ],
    widths=[5.0, 3.0, 4.0]
)
doc.add_paragraph()

bl('Yếu tố nhân khẩu học-xã hội liên quan đến lo âu:', bold=True)
tbl(
    ['Yếu tố', 'Tổng mẫu (N=450)', 'Có lo âu (n=108)', 'P'],
    [
        ['Nhóm tuổi 16\u201318', '239 (53,1%)', '66 (61,1%)', '0,032'],
        ['Nữ giới', '181 (40,2%)', '63 (58,3%)', '0,022'],
        ['Gia đình hạt nhân', '329 (73,1%)', '68 (63,0%)', '0,042'],
        ['Sống cùng gia đình', '354 (78,7%)', '68 (63,0%)', '0,035'],
        ['Mối quan hệ bạn bè tốt', '376 (83,6%)', '61 (56,5%)', '0,019'],
    ],
    widths=[4.5, 3.0, 3.0, 2.0]
)
doc.add_paragraph()

bl('Yếu tố lối sống liên quan đến lo âu:', bold=True)
tbl(
    ['Yếu tố lối sống', 'Có lo âu (%)', 'P', 'Ghi chú'],
    [
        ['Không hoạt động thể chất', '78,7% trong nhóm lo âu', '0,025', 'Có ý nghĩa'],
        ['Không hài lòng giấc ngủ', '27,8% trong nhóm lo âu', '0,018', 'Có ý nghĩa'],
        ['Hút thuốc lá (có)', '85,2% trong nhóm lo âu', '<0,001', 'Có ý nghĩa (92,4% mẫu hút thuốc — số liệu đáng ngờ)'],
        ['Giờ ngủ ít hơn BT', '63,9% trong nhóm lo âu', '0,072', 'Không ý nghĩa'],
        ['Sử dụng internet', '98,2% trong nhóm lo âu', '0,634', 'Không ý nghĩa'],
    ],
    widths=[4.0, 3.5, 2.0, 3.0]
)
doc.add_paragraph()

# BƯỚC 10: ĐỐI CHIẾU (KT3 — xác nhận 3 bước)
bh3('Khác biệt giới tính : Nữ giới luôn được báo cáo có tỷ lệ lo âu cao hơn nam giới trong y văn quốc tế (Lancet Psychiatry, 2020; Nag và cộng sự, 2019). Thực tế, trong nghiên cứu tại các bang Ấn Độ, nữ sinh có mức lo âu nặng hơn so với nam sinh (10,9% so với 3,8%, Nag và cộng sự, 2019). Chúng tôi xác nhận trong nhóm thanh thiếu niên tại trường đô thị tự trị này rằng nữ chiếm 58,3% trong nhóm có lo âu, trong khi nam chỉ chiếm 41,7% (P = 0,022) \u2014 phù hợp với xu hướng quốc tế.')

# BƯỚC 11+12: NHẬN XÉT + CẦU NỐI
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về yếu tố lối sống và mối quan hệ xã hội')
bl('*Hoạt động thể chất và giấc ngủ.* Nổi bật đối với nhóm học sinh có lo âu là tỷ lệ rất cao (78,7%) không hoạt động thể chất thường xuyên. Hoạt động thể chất dưới hình thức thể thao và trò chơi đã được chứng minh giúp cải thiện kỹ năng giải quyết vấn đề và giúp thanh thiếu niên đối phó tốt hơn với các vấn đề sức khỏe tâm thần (Soltanian và cộng sự, 2011). Đây là một yếu tố bảo vệ có thể can thiệp được, tương tự như phát hiện của Zhu và cộng sự (2025) rằng hoạt động ngoài trời 2\u20133 giờ/ngày giảm AOR = 0,557 cho trầm cảm.')
bl('*Mối quan hệ bạn bè.* Sự cảm nhận về mối quan hệ không tốt với bạn bè có liên quan có ý nghĩa thống kê với lo âu (P = 0,019): 43,5% học sinh có lo âu cảm nhận mối quan hệ bạn bè không tốt, so với 16,4% trong toàn mẫu. Phát hiện này tương tự với Ngo Anh Vinh và cộng sự (2024) tại Việt Nam, nơi mối quan hệ bạn bè kém cũng là yếu tố nguy cơ cho lo âu ở thanh thiếu niên dân tộc thiểu số.')
bl('*Nhóm tuổi vị thành niên muộn.* Nhóm tuổi 16\u201318 chiếm 61,1% trong nhóm lo âu (P = 0,032), gợi ý rằng giai đoạn vị thành niên muộn \u2014 với áp lực thi cử, lựa chọn nghề nghiệp và thay đổi nội tiết tố \u2014 là thời kỳ nguy cơ cao. Điều này phù hợp với nghiên cứu trên tạp chí Lancet cho thấy tỷ lệ rối loạn lo âu tăng nhanh ở thanh thiếu niên và thanh niên tại Ấn Độ.')

# BƯỚC 13: KẾT LUẬN (Kiến trúc B)
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy tỷ lệ lo âu 24% (108/450) ở thanh thiếu niên tại trường đô thị tự trị Ấn Độ \u2014 với mức nhẹ đến trung bình chiếm đa số (71,3%), gợi ý rằng cần chú ý đến phát hiện sớm ở giai đoạn vị thành niên muộn (16\u201318 tuổi) khi tỷ lệ lo âu cao nhất. Các yếu tố lối sống có thể can thiệp được \u2014 thiếu hoạt động thể chất, không hài lòng với giấc ngủ \u2014 là những mục tiêu tiềm năng cho chương trình phòng ngừa tại trường học. Trong bối cảnh trường đô thị tự trị như nghiên cứu hiện tại, nhu cầu triển khai sàng lọc sức khỏe tâm thần định kỳ và can thiệp dựa trên bằng chứng, đặc biệt là tăng cường hoạt động thể chất và hỗ trợ mối quan hệ bạn bè, là thiết yếu để giảm gánh nặng rối loạn lo âu ở thanh thiếu niên.', bold=True)

# PHẢN BIỆN
rh2('Phản biện')
bl('Tạp chí MedPulse International Journal of Psychology không có chỉ số tác động (impact factor) và không được lập chỉ mục trong PubMed, Scopus hay Web of Science \u2014 quy trình phản biện ngang hàng không rõ ràng, làm giảm độ tin cậy học thuật. Phương pháp chọn mẫu không rõ ràng: chỉ nêu "3 trường được xác định" mà không mô tả tiêu chí chọn hay quy trình lấy mẫu ngẫu nhiên (so với Saikia và cộng sự, 2023, chọn ngẫu nhiên 10/120 trường). Chỉ sử dụng chi-square (đơn biến) mà không dùng hồi quy logistic đa biến \u2014 không thể đánh giá ảnh hưởng độc lập của từng yếu tố sau khi kiểm soát biến gây nhiễu (Nakie và cộng sự, 2022, sử dụng AOR làm chuẩn tối thiểu). Dữ liệu hút thuốc đáng ngờ: 92,4% (416/450) học sinh được ghi nhận "có hút thuốc" \u2014 tỷ lệ này cao bất thường đối với học sinh trung học, gợi ý có thể có lỗi mã hóa dữ liệu hoặc lỗi dịch câu hỏi. Tuy nhiên, nghiên cứu có pilot test trên 30 học sinh và sử dụng GAD-7 \u2014 công cụ đã được xác thực quốc tế (Spitzer và cộng sự, 2006).')

# HƯỚNG NC
rh2('Hướng nghiên cứu tiếp theo')
bl('Tái lặp nghiên cứu với cỡ mẫu lớn hơn, chọn mẫu ngẫu nhiên có hệ thống và phân tích hồi quy đa biến. Xác minh và giải thích tỷ lệ hút thuốc bất thường (92,4%). So sánh trường đô thị tự trị với trường tư thục và trường nông thôn trong cùng khu vực. Nghiên cứu can thiệp tăng cường hoạt động thể chất tại trường để giảm lo âu \u2014 dựa trên bằng chứng 78,7% nhóm lo âu thiếu hoạt động thể chất. Xuất bản kết quả trên tạp chí có chỉ mục quốc tế để tăng khả năng so sánh liên nghiên cứu.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50 Trung bình-Thấp. Sử dụng GAD-7 đã xác thực quốc tế nhưng tạp chí không có impact factor, thiếu phân tích đa biến, dữ liệu hút thuốc đáng ngờ, và phương pháp chọn mẫu không rõ ràng.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('03_Mandaknalli_Malusare_2021.docx')
print('03_Mandaknalli_Malusare_2021.docx OK')
