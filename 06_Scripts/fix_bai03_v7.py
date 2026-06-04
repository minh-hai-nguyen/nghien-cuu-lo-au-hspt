# -*- coding: utf-8 -*-
"""
Bài 03 v6 -> v7: sửa dứt điểm Bảng 8 (model summary).
Lỗi: bảng có ô gộp (gridSpan=2) lệch nhau giữa hàng tiêu đề và hàng dữ liệu,
khiến lần rebuild v6 ghi đè nhầm -> mất cột "R² hiệu chỉnh" và giá trị SEE 3,638.
Cách sửa: gỡ toàn bộ gridSpan để mỗi hàng có đúng 6 ô, rồi đặt lại 6 + 6 giá trị.
Giá trị đúng (xác nhận từ bảng gốc v4/v5): R=0,177; R²=0,031; R² hiệu chỉnh=0,027;
SEE=3,638; Durbin-Watson=2,010.
"""
import io, sys, copy
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.oxml.ns import qn

IN = "bai-bao-khgdvn/bài-03/Công Thị Hằng_v6.docx"
OUT = "bai-bao-khgdvn/bài-03/Công Thị Hằng_v7.docx"

d = Document(IN)
tb = d.tables[7]  # Bảng 8


def set_cell(cell, text):
    p = cell.paragraphs[0]
    if p.runs:
        p.runs[0].text = text
        for r in p.runs[1:]:
            r.text = ""
    else:
        p.add_run(text)


def unmerge_row(tr):
    """Gỡ gridSpan: mỗi ô gộp -> tách thành nhiều ô đơn."""
    for tc in list(tr.findall(qn("w:tc"))):
        tcPr = tc.find(qn("w:tcPr"))
        if tcPr is None:
            continue
        gs = tcPr.find(qn("w:gridSpan"))
        if gs is not None:
            span = int(gs.get(qn("w:val")))
            tcPr.remove(gs)
            for _ in range(span - 1):
                ntc = copy.deepcopy(tc)
                tc.addnext(ntc)


hdr = ["Mô hình", "R", "R²", "R² hiệu chỉnh",
       "Sai số chuẩn của ước lượng", "Durbin – Watson"]
dat = ["1", "0,177", "0,031", "0,027", "3,638", "2,010"]

for ridx, vals in [(0, hdr), (1, dat)]:
    unmerge_row(tb.rows[ridx]._tr)
    cells = tb.rows[ridx].cells
    print(f"  Hàng {ridx}: {len(cells)} ô sau khi gỡ gộp")
    for ci in range(6):
        set_cell(cells[ci], vals[ci])

# kiểm tra
for ri in range(2):
    print(f"  R{ri}:", [c.text.strip() for c in tb.rows[ri].cells])

cp = d.core_properties
cp.author = ""
cp.last_modified_by = ""
d.save(OUT)
print(f"Đã lưu: {OUT}")
