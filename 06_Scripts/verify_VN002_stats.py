# -*- coding: utf-8 -*-
"""Verify VN002 FULL — all critical stats vs PDF gốc (meta-review vòng 2)."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

PDF_TXT = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_pdf_full.txt'
with open(PDF_TXT, encoding='utf-8') as f:
    pdf = f.read()
pdf_norm = re.sub(r'\s+', ' ', pdf).lower()

from docx import Document
docx = Document('03_Ban-dich/VN002_VNAMHS_2022_National_FULL.docx')
dtext = '\n'.join(p.text for p in docx.paragraphs if p.text.strip())
for t in docx.tables:
    for r in t.rows:
        for c in r.cells:
            dtext += ' ' + c.text
dtext_norm = re.sub(r'\s+', ' ', dtext)

# Claims to verify — each must be supported by variants in PDF AND present in translation
claims = [
    ('5.996 cặp', ['5,996', '5996', '5.996'], ['5.996', '5996']),
    ('38 tỉnh', ['38 provinces', '38 selected'], ['38 tỉnh']),
    ('7.599 hộ tiếp cận', ['7,599', '7599'], ['7.599', '7599']),
    ('81,1 % response', ['81.1%', '81.1 %', '81,1'], ['81,1']),
    ('6.048 hộ đủ điều kiện', ['6,048', '6048'], ['6.048', '6048']),
    ('Tuổi TB 13,3', ['13.3', '13,3'], ['13,3']),
    ('54,2 % nhóm 10-13', ['54.2', '54,2'], ['54,2']),
    ('94,5 % đang đi học', ['94.5', '94,5'], ['94,5']),
    ('94,6 % chưa làm việc', ['94.6', '94,6'], ['94,6']),
    ('21,7 % vấn đề SKTT', ['21.7', '21,7'], ['21,7']),
    ('3,3 % rối loạn', ['3.3%', '3,3'], ['3,3']),
    ('18,6 % lo âu', ['18.6', '18,6'], ['18,6']),
    ('4,3 % trầm cảm', ['4.3', '4,3'], ['4,3']),
    ('2,8 % ADHD', ['2.8', '2,8'], ['2,8']),
    ('1,0 % PTSD', ['1.0%', '1,0'], ['1,0']),
    ('0,7 % conduct', ['0.7%', '0,7'], ['0,7']),
    ('2,3 % anxiety disorder', ['2.3%', '2,3'], ['2,3']),
    ('0,9 % MDD', ['0.9%', '0,9'], ['0,9']),
    ('0,5 % ADHD disorder', ['0.5%', '0,5'], ['0,5']),
    ('67,0 % impairment family', ['67.0', '67,0'], ['67,0']),
    ('47,0 % impairment peer', ['47.0', '47,0'], ['47,0']),
    ('45,4 % impairment school', ['45.4', '45,4'], ['45,4']),
    ('34,6 % impairment personal', ['34.6', '34,6'], ['34,6']),
    ('1,4 % ý nghĩ tự sát', ['1.4%', '1,4'], ['1,4']),
    ('0,4 % kế hoạch', ['0.4%', '0,4'], ['0,4']),
    ('0,2 % toan tự sát', ['0.2%', '0,2'], ['0,2']),
    ('1,6 % ever toan', ['1.6', '1,6'], ['1,6']),
    ('4,7 % ever self-harm', ['4.7', '4,7'], ['4,7']),
    ('8,4 % service use (có SKTT)', ['8.4', '8,4'], ['8,4']),
    ('6,5 % service use (all)', ['6.5', '6,5'], ['6,5']),
    ('56,2 % doctor/nurse', ['56.2', '56,2'], ['56,2']),
    ('5,1 % parent need', ['5.1', '5,1'], ['5,1']),
    ('37,7 % received all help', ['37.7', '37,7'], ['37,7']),
    ('73,9 % talk to family', ['73.9', '73,9'], ['73,9']),
    ('38,2 % talk to friend', ['38.2', '38,2'], ['38,2']),
    ('30,7 % exercise self-help', ['30.7', '30,7'], ['30,7']),
    ('7,7 % COVID problem', ['7.7', '7,7'], ['7,7']),
    ('71,5 % hộ giảm thu nhập', ['71.5', '71,5'], ['71,5']),
    ('12,3 % thiếu tiền nhu yếu', ['12.3', '12,3'], ['12,3']),
    ('80,3 % không tiếp cận dịch vụ', ['80.3', '80,3'], ['80,3']),
    ('69,2 % sợ COVID', ['69.2', '69,2'], ['69,2']),
    ('127 phỏng vấn viên', ['127 interviewers', '127 trained'], ['127 phỏng vấn']),
]

print(f'{"Claim":50s} {"In PDF":>8s} {"In VN002":>10s}')
print('=' * 72)
ok = 0; fail = 0
for label, pdf_variants, vn_variants in claims:
    pdf_hit = any(v.lower() in pdf_norm for v in pdf_variants)
    vn_hit = any(v in dtext_norm for v in vn_variants)
    status = '✓' if pdf_hit and vn_hit else ('⚠' if pdf_hit else '✗')
    if pdf_hit and vn_hit:
        ok += 1
    else:
        fail += 1
    print(f'{label:50s} {"✓" if pdf_hit else "✗":>8s} {"✓" if vn_hit else "✗":>10s}  {status}')

print('=' * 72)
print(f'PASS: {ok}/{len(claims)} | FAIL: {fail}')
print(f'Meta-review vòng 2 (stats vs PDF + translation): {100*ok/len(claims):.1f}% verified')
