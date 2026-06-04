# -*- coding: utf-8 -*-
"""Sinh ban TONG QUAN RUT GON ~30 trang (~12.000 tu) theo y kien thay.
Strategy: copy FULL v1 -> remove select paragraphs (peripheral) -> expand Can thiep.
26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_TheoChuDe_FULL_v1_26052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_TheoChuDe_30trang_v2_26052026.docx')

# ============================================================
# DOC SETUP
# ============================================================
def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.5)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)

def P(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    return p

# ============================================================
# DOC PARAGRAPHS TO REMOVE (key phrases for content match)
# ============================================================
# GIU LAI:
#  - Tat ca doan "khoang trong VN" / "Nhan xet" (de trich sang Ly do + Kien nghi)
#  - Cac doan cuoi cua moi section (gap analysis)
# CAT:
#  - Cac doan peripheral o GIUA section (chi tiet quoc te, dinh nghia phai sinh)
REMOVE_KEYWORDS = [
    # Section 1: Mở rộng so sánh quốc tế Nakie/Bhardwaj/Alharbi (~400 từ - peripheral QT)
    'Mở rộng so sánh quốc tế cho thấy biên độ tỷ lệ tương tự',
    # Section 1: Một nghiên cứu so sánh quốc tế Jefferies (~300 từ - peripheral QT)
    'Một nghiên cứu so sánh quốc tế đặc biệt đáng lưu ý cho bối cảnh Việt Nam',
    # Section 2: Phân tích dữ liệu Trung Quốc Xu detail (~300 từ - peripheral data Xu)
    'Phân tích dữ liệu Trung Quốc của Xu và cộng sự (2021) trên 373.216 học sinh',
    # Section 2: Chen detail (~250 từ - peripheral)
    'Phân tích của Chen và cộng sự (2023) trên 63.205 học sinh tại Tứ Xuyên',
    # Section 3: Tiêu chí chẩn đoán DSM-5 chi tiết (~250 từ - definition detail)
    'Tiêu chí chẩn đoán DSM-5 cho rối loạn lo âu lan tỏa yêu cầu',
    # Section 3: Đồng diễn 3 dạng (~400 từ - methodological)
    'Tỷ lệ đồng diễn (comorbidity) giữa ba dạng rối loạn lo âu',
    # Section 4: Tương tác giữa các nhóm (~350 từ - analytical, ref Bai 1 da co)
    'Tương tác giữa các nhóm yếu tố nguy cơ là một chủ đề',
    # Section 4: Mở rộng phân tích áp lực học tập (~300 từ - peripheral OECD data)
    'Mở rộng phân tích về áp lực học tập như yếu tố nguy cơ',
    # Section 5: Cơ chế bảo vệ gắn bó trường học (~300 từ - peripheral mechanism)
    'Cơ chế bảo vệ của gắn bó trường học hoạt động qua ba kênh',
    # Section 1: Tổng hợp xu hướng (~400 từ - peripheral, just summary of what said above)
    'Tổng hợp các nghiên cứu Việt Nam giai đoạn 2020-2026 cho thấy ba xu hướng',
    # Section 3: lịch sử DSM-IV (~250 từ - historical detail)
    'Rối loạn lo âu chia ly có lịch sử thay đổi tiêu chí chẩn đoán',
    # Section 5.6: Văn hóa Việt Nam đặc thù (~280 từ - peripheral mechanism)
    'Đối với học sinh trung học cơ sở Việt Nam, đặc thù văn hóa cộng đồng',
    # Section 5: KHONG CAT khoang trong YTBV (di vao Kien nghi)
    # Section 7: KHONG CAT "Đặt trong bối cảnh quốc tế" (gap analysis cua LA - di vao Ly do)
]

# ============================================================
# NEW PARAGRAPHS TO ADD TO CAN THIEP (1.1.6) ~+1200 tu
# ============================================================
# Will be inserted before the last paragraph of section 1.1.6
NEW_CT_PARAGRAPHS = [
    'Một nhân tố quan trọng cho việc lựa chọn mô hình can thiệp tại Việt Nam là điều kiện thực tế của hệ thống chăm sóc sức khỏe tâm thần học đường. Theo các báo cáo nội bộ của Bộ Giáo dục và Đào tạo, đa số trường trung học cơ sở Việt Nam hiện đã có tổ tư vấn tâm lý theo Thông tư 31/2017/TT-BGDĐT, song phần lớn hoạt động tập trung vào tham vấn cá nhân khi học sinh chủ động tìm đến, chưa có chương trình phòng ngừa có cấu trúc. Mặt khác, nhân lực tâm lý lâm sàng được đào tạo bài bản về can thiệp nhận thức — hành vi cho trẻ em — vị thành niên còn rất hạn chế, chủ yếu tập trung tại các bệnh viện chuyên khoa ở Hà Nội và Thành phố Hồ Chí Minh. Tỷ lệ học sinh có thể tiếp cận dịch vụ chuyên khoa khi cần — kể cả ở các thành phố lớn — vẫn ở mức rất thấp. Bối cảnh này gợi ý mô hình can thiệp ưu tiên cho Việt Nam là can thiệp tại trường học (kênh hai) bổ trợ bằng nền tảng số (kênh ba), với can thiệp lâm sàng (kênh một) dành cho các trường hợp triệu chứng nặng có yêu cầu chuyển tuyến chuyên khoa.',

    'Một câu hỏi thiết kế quan trọng cho chương trình tập huấn phòng ngừa cho học sinh trung học cơ sở Việt Nam là việc lựa chọn giữa mô hình phổ quát (universal) và mô hình chọn lọc (selective/indicated). Mô hình phổ quát có ưu điểm là không gắn nhãn học sinh "có vấn đề", giảm kỳ thị và tận dụng được khung môn học chính khóa — đặc biệt môn Hoạt động trải nghiệm — Hướng nghiệp của Chương trình giáo dục phổ thông 2018. Mô hình chọn lọc và chỉ định có cỡ hiệu ứng cao hơn nhưng đòi hỏi quy trình sàng lọc, gắn nhãn và có thể tạo cảm giác bị phân biệt cho học sinh được chọn. Phân tích tổng hợp của Jagiello và cộng sự (2025) trên 31 nghiên cứu can thiệp áp lực học tập tại trường từ 13 quốc gia cho thấy cả hai mô hình đều có hiệu quả, song với cỡ hiệu ứng khác nhau: mô hình phổ quát có cỡ hiệu ứng nhỏ-vừa nhưng tiếp cận được nhiều học sinh hơn, mô hình chọn lọc có cỡ hiệu ứng lớn hơn nhưng phạm vi hạn chế. Một thiết kế hai tầng — tầng phổ quát 4-6 buổi tích hợp môn học, tầng chọn lọc 8-12 buổi cho học sinh được sàng lọc — là cách kết hợp ưu điểm của cả hai mô hình, đồng thời phù hợp với cấu trúc thực tế của trường trung học cơ sở Việt Nam.',

    'Vai trò của giáo viên chủ nhiệm và giáo viên tâm lý học đường trong triển khai chương trình can thiệp tại trường là một chủ đề cần được làm rõ. Khác với chuyên gia tâm lý lâm sàng, giáo viên chủ nhiệm có lợi thế về tần suất tiếp xúc hằng ngày với học sinh và hiểu biết về hoàn cảnh gia đình, song lại hạn chế về thời gian và năng lực can thiệp lâm sàng. Mô hình "giáo viên chủ nhiệm là điều phối viên — chuyên gia tâm lý là người giám sát" đã được áp dụng tại một số nước châu Á có thể là phương án phù hợp cho điều kiện Việt Nam khi nguồn lực chuyên môn còn hạn chế. Cụ thể, giáo viên chủ nhiệm được đào tạo về sàng lọc sơ cấp và kỹ thuật can thiệp cơ bản trong khuôn khổ chương trình môn học, trong khi nhà tâm lý đảm nhận giám sát chuyên môn, tư vấn ca khó và tiếp nhận các trường hợp cần can thiệp sâu hơn. Đánh giá hệ thống của Jagiello và cộng sự (2025) ghi nhận một số ít chương trình do giáo viên dẫn dắt vẫn đạt hiệu quả khi có đào tạo và giám sát chuyên môn đầy đủ — phát hiện có ý nghĩa thực tế cho Việt Nam.',
]

# ============================================================
# BUILD V2
# ============================================================
src = Document(SRC)
new = doc_init()

# Track: section currently being processed
n_removed = 0
n_added_ct = 0
in_ct_section = False
ct_paragraphs_added = False  # to ensure we only add NEW paragraphs once

def should_remove(text):
    """Check if paragraph text matches any REMOVE_KEYWORDS."""
    for kw in REMOVE_KEYWORDS:
        if text.startswith(kw[:50]):
            return True
    return False

def copy_para(src_p):
    """Copy paragraph from src to new, preserving format."""
    style = src_p.style.name
    if style == 'List Bullet':
        new_p = new.add_paragraph(style='List Bullet')
    elif style.startswith('Heading'):
        level = int(style[-1]) if style[-1].isdigit() else 2
        new_p = new.add_heading('', level=level)
    else:
        new_p = new.add_paragraph()
        new_p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        src_fli = src_p.paragraph_format.first_line_indent
        if src_fli and src_fli.cm > 0:
            new_p.paragraph_format.first_line_indent = src_fli
        src_li = src_p.paragraph_format.left_indent
        if src_li and src_li.cm > 0:
            new_p.paragraph_format.left_indent = src_li
        src_ls = src_p.paragraph_format.line_spacing
        if src_ls:
            new_p.paragraph_format.line_spacing = src_ls

    # Copy runs
    for src_r in src_p.runs:
        if not src_r.text:
            continue
        new_r = new_p.add_run(src_r.text)
        new_r.bold = src_r.bold
        new_r.italic = src_r.italic
        new_r.underline = src_r.underline
        font_name = src_r.font.name or 'Times New Roman'
        new_r.font.name = font_name
        font_size = src_r.font.size
        new_r.font.size = font_size if font_size else Pt(13)
        if src_r.font.color and src_r.font.color.rgb:
            try:
                new_r.font.color.rgb = src_r.font.color.rgb
            except Exception:
                pass

    # Heading fix: in case heading - need to rebuild text
    if style.startswith('Heading'):
        # Add heading text properly
        for r in new_p.runs:
            r.font.name = 'Times New Roman'
            r.font.color.rgb = RGBColor(0, 0, 0)

    return new_p


# Process each src paragraph
for src_p in src.paragraphs:
    text = src_p.text.strip()

    # Track section
    if src_p.style.name == 'Heading 3':
        if text.startswith('1.1.6'):
            in_ct_section = True
            ct_paragraphs_added = False
        elif text.startswith('1.1.7'):
            # Before moving out of CT section, add new CT paragraphs
            if in_ct_section and not ct_paragraphs_added:
                for npar in NEW_CT_PARAGRAPHS:
                    P(new, npar, indent=True)
                    n_added_ct += 1
                ct_paragraphs_added = True
            in_ct_section = False

    # Skip if matches remove keyword
    if text and should_remove(text):
        n_removed += 1
        continue

    # Copy the paragraph
    copy_para(src_p)

# Edge case: if 1.1.7 was the last heading, CT paragraphs might not be added.
# In our structure, 1.1.7 exists after 1.1.6, so the loop handles it.

new.save(OUT)

# Verify
words = sum(len(p.text.split()) for p in new.paragraphs)
paras = sum(1 for p in new.paragraphs if p.text.strip())
print(f"Da luu: {OUT}")
print(f"Words: {words}")
print(f"Paragraphs: {paras}")
print(f"Removed: {n_removed} paragraphs")
print(f"Added to CT: {n_added_ct} paragraphs")
pages_400 = words / 400
print(f"Est pages: {pages_400:.1f}")
