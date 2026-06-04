# -*- coding: utf-8 -*-
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Times New Roman'
font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3)
    section.right_margin = Cm(2)

def add_h(text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = 'Times New Roman'
        run.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_b(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    return p

def add_p(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    return p

def add_i(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    return p

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def tbl(headers, rows, widths=None):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    if widths:
        for i, w in enumerate(widths):
            for row in t.rows:
                row.cells[i].width = Cm(w)
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    return t

def info_tbl(data):
    """Bảng thông tin 2 cột: Mục | Chi tiết, widths 3.0 + 13.0"""
    t = doc.add_table(rows=len(data), cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        row.cells[0].width = Cm(3.0)
        row.cells[1].width = Cm(13.0)
    for i, (k, v) in enumerate(data):
        c0 = t.rows[i].cells[0]
        c0.text = k
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        shade(c0, 'D9E2F3')
        c1 = t.rows[i].cells[1]
        c1.text = v
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
    return t

# ==================== TITLE ====================
title = doc.add_heading('', level=0)
run = title.add_run('BÁO CÁO TỔNG HỢP\n11 BÀI NGHIÊN CỨU VỀ LO ÂU Ở HỌC SINH THCS VÀ THPT\n(Việt Nam – Đông Nam Á – Thế giới, 2024–2026)')
run.font.name = 'Times New Roman'
run.font.size = Pt(16)
run.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_i('Ngày báo cáo: 27/03/2026. Nguồn: các bài nghiên cứu từ 2024 đến nay trên PubMed, Frontiers, PLOS ONE, Springer, ScienceDirect.')
doc.add_paragraph()

# ==================== PHẦN A: VIỆT NAM ====================
add_h('PHẦN A: CÁC NGHIÊN CỨU TẠI VIỆT NAM', 1)

# --- BÀI 1 ---
add_h('Bài 1: Phạm Thị Thu Hoa và cs. (2024) — Lo âu và chiến lược ứng phó ở HS THPT Hà Nội sau COVID-19', 2)
info_tbl([
    ('Tác giả', 'Phạm Thị Thu Hoa, Đỗ Thị Trang, Nguyễn Thị Liên, Ngô Anh Vinh'),
    ('Tạp chí', 'Frontiers in Public Health (IF ≈ 5.2)'),
    ('Năm', '2024'),
    ('Mẫu', '3.910 học sinh THPT tại 13 trường ở Hà Nội'),
    ('Công cụ', 'GAD-7 (Cronbach α = 0,916); phỏng vấn sâu 20 HS'),
    ('Thiết kế', 'Nghiên cứu hỗn hợp (cắt ngang + định tính)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu sử dụng phương pháp hỗn hợp, kết hợp khảo sát cắt ngang trên 3.910 học sinh THPT tại 13 trường ở Hà Nội (tháng 10–11/2021) với phỏng vấn sâu qua điện thoại cho 20 học sinh. Công cụ đo lo âu là thang GAD-7 với 7 câu hỏi, điểm 0–21. Thang đạt độ tin cậy cao với Cronbach α = 0,916. Dữ liệu phân tích bằng SPSS 20.0, mức ý nghĩa p ≤ 0,05.')

add_b('Kết quả chính')
add_p('Tỷ lệ học sinh có triệu chứng lo âu tổng thể đạt 40,6%. Trong đó, lo âu nhẹ chiếm 23,9%, trung bình 10,9%, và nặng 5,8%. Triệu chứng phổ biến nhất là "dễ cáu gắt, khó chịu" (mean = 1,01), tiếp theo là "lo lắng quá mức về nhiều vấn đề" (mean = 0,96) và "cảm thấy bồn chồn, lo lắng" (mean = 0,89).')
add_p('Có sự khác biệt có ý nghĩa thống kê theo giới tính, vị trí địa lý, và khối lớp (p < 0,05). Phân tích định tính cho thấy các nguồn lo âu chính gồm: áp lực học tập, mối quan hệ xã hội, kỳ thị từ bạn bè, và kỳ vọng gia đình.')

doc.add_paragraph()

# --- BÀI 2 ---
add_h('Bài 2: Ngô Anh Vinh và cs. (2024) — SKTT ở HS dân tộc thiểu số tại Lạng Sơn', 2)
info_tbl([
    ('Tác giả', 'Ngô Anh Vinh, Vũ Thị Mỹ Hạnh, Đỗ Thị Bích Vân, Dương Anh Tài, Lê Thị Thanh Thùy'),
    ('Tạp chí', 'Journal of Affective Disorders Reports (Elsevier)'),
    ('Năm', '2024 (Vol. 17)'),
    ('Mẫu', '845 học sinh nội trú dân tộc thiểu số tại Lạng Sơn'),
    ('Công cụ', 'DASS-21; Bảng hỏi ACEs (trải nghiệm bất lợi thời thơ ấu)'),
    ('Thiết kế', 'Nghiên cứu cắt ngang')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu cắt ngang trên 845 học sinh nội trú tại các trường dân tộc thiểu số ở Lạng Sơn. Sử dụng thang DASS-21 đo trầm cảm, lo âu, căng thẳng và Bảng hỏi ACEs đo trải nghiệm bất lợi thời thơ ấu. Phân tích bằng hồi quy logistic đa biến và hồi quy Tobit.')

add_b('Kết quả chính')
add_p('Tỷ lệ vấn đề sức khỏe tâm thần rất cao: lo âu 54,4%, trầm cảm 59,0%, căng thẳng 24,7%. Có 48,9% học sinh từng trải qua ít nhất 1 trải nghiệm bất lợi thời thơ ấu (ACEs), trung bình 1,1 ACEs. Số lượng ACEs tương quan dương với điểm lo âu (hệ số = 0,28), căng thẳng (0,28) và trầm cảm (0,16).')
add_p('Mối quan hệ bạn bè kém có liên quan đáng kể đến lo âu, căng thẳng và trầm cảm gia tăng. Sử dụng internet hàng ngày hoặc không sử dụng có điểm lo âu thấp hơn so với nhóm chỉ dùng cuối tuần. Nguồn thông tin từ bạn bè và truyền hình có mối liên hệ nghịch với căng thẳng và lo âu.')

doc.add_paragraph()

# --- BÀI 3 ---
add_h('Bài 3: Trần Thảo Vi và cs. (2024) — Căng thẳng học tập ở HS THCS tại Huế: nghiên cứu dọc 3 năm', 2)
info_tbl([
    ('Tác giả', 'Trần Thảo Vi (ĐH Y khoa Tokyo & ĐH Y Dược Huế) và cộng sự'),
    ('Tạp chí', 'Journal of Rural Medicine (IF ≈ 1.0)'),
    ('Năm', '2024 (Vol. 19, No. 4, pp. 279–290)'),
    ('Mẫu', '611 HS lớp 6 → theo dõi đến lớp 9 (341 HS hoàn thành); 4 trường THCS tại Huế'),
    ('Công cụ', 'ESSA (Educational Stress Scale for Adolescents) – 16 câu, 5 miền'),
    ('Thiết kế', 'Nghiên cứu dọc (longitudinal cohort) 2018–2021')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu dọc 3 năm (2018–2021) theo dõi 611 học sinh lớp 6 tại 4 trường THCS ở Huế đến lớp 9 (341 HS hoàn thành). Sử dụng thang ESSA gồm 16 câu hỏi thuộc 5 miền: áp lực học tập, lo lắng về điểm, chán nản, kỳ vọng bản thân, và khối lượng công việc. Phân tích bằng hồi quy tuyến tính.')

add_b('Kết quả chính')
add_p('Điểm ESSA tăng đáng kể từ 46,4 ± 7,6 (2018) lên 53,5 ± 10,8 (2021), cho thấy căng thẳng học tập gia tăng rõ rệt theo thời gian. Giới tính nữ có điểm cao hơn nam (β = −2,85 cho nam). Số anh chị em nhiều hơn (β = 2,24), trình độ học vấn cha cao (β = 3,20), điểm học thấp (β = −1,79), và học thêm (β = 4,73) đều liên quan đến căng thẳng cao hơn.')
add_p('Ngủ đủ 8 giờ trở lên có tác dụng bảo vệ, làm giảm điểm "chán nản". Hoạt động thể chất không có mối liên hệ có ý nghĩa. Thời điểm khảo sát lần cuối trùng với đại dịch COVID-19, có thể làm tăng thêm căng thẳng cho HS chuẩn bị thi chuyển cấp.')

doc.add_paragraph()

# --- BÀI 4 ---
add_h('Bài 4: Nguyễn Ngọc Bảo Quyên và cs. (2025) — Trầm cảm, lo âu, căng thẳng ở HS THPT Hà Nội', 2)
info_tbl([
    ('Tác giả', 'Nguyễn Ngọc Bảo Quyên và cộng sự'),
    ('Tạp chí', 'Tạp chí Y học Cộng đồng (Vol. 66)'),
    ('Năm', '2025'),
    ('Mẫu', '501 học sinh THPT tại Hà Nội'),
    ('Công cụ', 'DASS-21'),
    ('Thiết kế', 'Nghiên cứu cắt ngang (tháng 8–12/2023)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu cắt ngang trên 501 học sinh THPT tại Hà Nội, tiến hành từ tháng 8 đến 12/2023. Sử dụng thang DASS-21 (21 câu hỏi) qua bảng khảo sát trực tuyến. Phân tích mối liên quan giữa giới tính, khối lớp, loại trường với các chỉ số SKTT.')

add_b('Kết quả chính')
add_p('Tỷ lệ rất cao: 86,2% HS có dấu hiệu lo âu, 78,8% trầm cảm, 76,6% căng thẳng. Các trường hợp nặng và rất nặng chiếm tỷ lệ cao hơn đáng kể ở nữ so với nam (p < 0,05). Không tìm thấy mối liên hệ có ý nghĩa giữa khối lớp hoặc loại trường với các vấn đề SKTT.')
add_p('Mặc dù tỷ lệ rất cao, cần lưu ý DASS-21 là công cụ sàng lọc chứ không phải chẩn đoán lâm sàng. Tỷ lệ cao có thể phản ánh phần nào phương pháp khảo sát trực tuyến và bối cảnh hậu COVID-19.')

doc.add_paragraph()

# --- BÀI 5 ---
add_h('Bài 5: Hoàng Trung Học & Nguyễn Thùy Dung (2025) — Lo âu, trầm cảm ở VTN Việt Nam trong và sau COVID-19', 2)
info_tbl([
    ('Tác giả', 'Hoàng Trung Học (Học viện QLGD) và Nguyễn Thùy Dung (Viện NC Tâm lý Ứng dụng)'),
    ('Tạp chí', 'American Journal of Psychiatric Rehabilitation (Vol. 28, No. 1)'),
    ('Năm', '2025'),
    ('Mẫu', '8.473 vị thành niên tại 6 tỉnh/thành Việt Nam'),
    ('Công cụ', 'DASS-21 phiên bản tiếng Việt'),
    ('Thiết kế', 'Nghiên cứu cắt ngang, so sánh hai giai đoạn (trong và sau COVID-19)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu cắt ngang quy mô lớn trên 8.473 vị thành niên tại 6 tỉnh/thành Việt Nam. Sử dụng DASS-21 phiên bản tiếng Việt. So sánh hai thời điểm: trong đại dịch COVID-19 và sau đại dịch, nhằm đánh giá xu hướng thay đổi các chỉ số SKTT.')

add_b('Kết quả chính')
add_p('Trong đại dịch: căng thẳng 65,5%, lo âu 41,5%, trầm cảm 34,2%. Sau đại dịch: căng thẳng 55,4%, lo âu 25,4%, trầm cảm 20,1%. Tất cả các chỉ số đều giảm sau đại dịch nhưng vẫn ở mức cao.')

tbl(['Chỉ số', 'Trong đại dịch', 'Sau đại dịch', 'Thay đổi'],
    [['Căng thẳng', '65,5%', '55,4%', '↓ 10,1%'],
     ['Lo âu', '41,5%', '25,4%', '↓ 16,1%'],
     ['Trầm cảm', '34,2%', '20,1%', '↓ 14,1%']],
    widths=[4.0, 3.5, 3.5, 3.0])
doc.add_paragraph()

add_p('Các yếu tố nguy cơ: chất lượng mối quan hệ cha mẹ–con kém, sử dụng thiết bị điện tử quá mức, chất lượng giấc ngủ kém. Yếu tố bảo vệ: mối quan hệ gia đình tích cực (đặc biệt chống lo âu), hoạt động ngoại khóa và lối sống lành mạnh (sau đại dịch).')

doc.add_paragraph()

# ==================== PHẦN B: ĐÔNG NAM Á ====================
add_h('PHẦN B: CÁC NGHIÊN CỨU TẠI ĐÔNG NAM Á', 1)

# --- BÀI 6 ---
add_h('Bài 6: Sarfika và cs. (2025) — Yếu tố dự báo và bảo vệ SKTT ở VTN Indonesia', 2)
info_tbl([
    ('Tác giả', 'Rika Sarfika, I Made Moh. Yanuar Saifudin, Anggi Lukman Wicaksana và cộng sự'),
    ('Tạp chí', 'International Journal of Social Psychiatry'),
    ('Năm', '2025'),
    ('Mẫu', 'Nghiên cứu đa trung tâm tại Indonesia (vị thành niên)'),
    ('Công cụ', 'Dựa trên lý thuyết nhận thức xã hội (SCT)'),
    ('Thiết kế', 'Nghiên cứu cắt ngang đa trung tâm')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu cắt ngang đa trung tâm tại Indonesia, dựa trên khung lý thuyết nhận thức xã hội (Social Cognitive Theory – SCT). Đánh giá các yếu tố dự báo và yếu tố bảo vệ đối với trầm cảm, lo âu và căng thẳng ở vị thành niên Indonesia.')

add_b('Kết quả chính')
add_p('Mô hình hồi quy giải thích 57,1% phương sai trầm cảm, 42,8% lo âu, và 46,3% căng thẳng. Giới tính, áp lực học tập, và áp lực bạn bè là các yếu tố dự báo chính của rối loạn tâm thần. Các yếu tố bảo vệ bao gồm: lòng tự trọng, khả năng phục hồi, hỗ trợ xã hội và hoạt động thể chất.')
add_p('Kết quả cho thấy cần xây dựng các chương trình can thiệp tại trường học tập trung vào nâng cao lòng tự trọng, phát triển khả năng phục hồi, và tăng cường hỗ trợ xã hội cho vị thành niên Indonesia.')

doc.add_paragraph()

# --- BÀI 7 ---
add_h('Bài 7: Ayeras và cs. (2024) — Lo âu xã hội ở HS THPT Philippines', 2)
info_tbl([
    ('Tác giả', 'Sophia Randa F. Ayeras, Johnrev B. Bumanlag, Blessy Faith D. De Guzman, Belle Marienne B. Reyes, Anna Klara F. Ruiz, Joseph A. Villarama'),
    ('Tạp chí', 'Journal of Interdisciplinary Perspectives (Vol. 2, No. 7)'),
    ('Năm', '2024'),
    ('Mẫu', '12 học sinh THPT tại Central Luzon State University, Philippines'),
    ('Công cụ', 'Phỏng vấn bán cấu trúc, phân tích chủ đề'),
    ('Thiết kế', 'Nghiên cứu hiện tượng học (phenomenological qualitative study)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu định tính theo phương pháp hiện tượng học, dựa trên lý thuyết nhận thức xã hội. Phỏng vấn bán cấu trúc 12 học sinh THPT có lo âu xã hội tại Central Luzon State University, Philippines. Dữ liệu ghi âm và phân tích chủ đề (thematic analysis).')

add_b('Kết quả chính')
add_p('Nghiên cứu xác định nhiều yếu tố kích hoạt lo âu xã hội ở vị thành niên: môi trường, trải nghiệm tiêu cực, và xung đột nội tâm. Học sinh có lo âu xã hội gặp khó khăn trong học tập và mối quan hệ bạn bè. Lo âu xã hội ảnh hưởng tiêu cực đến lòng tự trọng.')
add_p('Chiến lược ứng phó hiệu quả giúp giảm lo âu, nhưng các chương trình hướng dẫn hiện tại tại trường còn thiếu. Tác giả khuyến nghị xây dựng các chương trình can thiệp dựa trên trường học, phát triển năng lực xã hội và hỗ trợ dài hạn cho HS.')

doc.add_paragraph()

# ==================== PHẦN C: THẾ GIỚI ====================
add_h('PHẦN C: CÁC NGHIÊN CỨU QUỐC TẾ', 1)

# --- BÀI 8 ---
add_h('Bài 8: Islam và cs. (2026) — Lo âu ở VTN đi học tại 59 quốc gia (GSHS)', 2)
info_tbl([
    ('Tác giả', 'Md. Amirul Islam và cộng sự'),
    ('Tạp chí', 'Journal of Affective Disorders (IF ≈ 6.6)'),
    ('Năm', '2026'),
    ('Mẫu', '179.937 vị thành niên 11–17 tuổi tại 59 quốc gia'),
    ('Công cụ', 'Global School-based Student Health Survey (GSHS)'),
    ('Thiết kế', 'Phân tích dữ liệu thứ cấp từ GSHS')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Phân tích dữ liệu thứ cấp từ Khảo sát Sức khỏe Học sinh Toàn cầu (GSHS) trên 179.937 vị thành niên 11–17 tuổi tại 59 quốc gia. Đây là một trong những nghiên cứu quy mô lớn nhất về lo âu ở VTN đi học, cho phép so sánh giữa các khu vực trên thế giới.')

add_b('Kết quả chính')
add_p('Tỷ lệ lo âu dao động lớn giữa các khu vực. Đông Địa Trung Hải có tỷ lệ cao nhất (14,61%, 95% CI: 14,59–14,63). Đông Nam Á có tỷ lệ thấp nhất (3,78%, 95% CI: 3,72–3,83). Khu vực châu Phi khoảng 11,4%, cao gấp gần 4 lần Đông Nam Á.')

tbl(['Khu vực', 'Tỷ lệ lo âu', '95% CI'],
    [['Đông Địa Trung Hải', '14,61%', '14,59–14,63'],
     ['Châu Phi', '~11,4%', '—'],
     ['Châu Âu', '~7–8%', '—'],
     ['Tây Thái Bình Dương', '~5–6%', '—'],
     ['Đông Nam Á', '3,78%', '3,72–3,83']],
    widths=[5.0, 4.0, 5.0])
doc.add_paragraph()

add_p('Các yếu tố nguy cơ: thiếu an ninh lương thực, giới nữ, tuổi lớn hơn, bị bắt nạt, thiếu bạn thân, cha mẹ thiếu thấu hiểu, lối sống ít vận động, và ý tưởng tự tử.')

doc.add_paragraph()

# --- BÀI 9 ---
add_h('Bài 9: Dong, Wang & Lin (2025) — Trầm cảm, lo âu, căng thẳng ở HS trung học tại Trung Quốc', 2)
info_tbl([
    ('Tác giả', 'Tingting Dong, Yumei Wang, Yuanjie Lin'),
    ('Tạp chí', 'PLOS ONE'),
    ('Năm', '2025'),
    ('Mẫu', '2.716 học sinh trung học tại Ya\'an, Tây Nam Trung Quốc'),
    ('Công cụ', 'DASS-21'),
    ('Thiết kế', 'Nghiên cứu cắt ngang (tháng 3–6/2022)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Nghiên cứu cắt ngang trên 2.716 HS trung học (45,3% nam, 54,7% nữ) tại Ya\'an, Tây Nam Trung Quốc, tiến hành tháng 3–6/2022. Sử dụng DASS-21 với thang 4 điểm (0–3). Tỷ lệ hoàn thành 93,66%. Phân tích bằng hồi quy logistic nhị phân, chi-square và t-test.')

add_b('Kết quả chính')
add_p('Tỷ lệ lo âu cao nhất: 41,4% (nhẹ 14,62%, trung bình 18,37%, nặng 4,09%, rất nặng 4,34%). Trầm cảm 24,4%. Căng thẳng 15,6%. HS THPT có tỷ lệ cao hơn HS THCS ở cả ba chỉ số (p < 0,001).')

tbl(['Chỉ số', 'Tổng', 'Nhẹ', 'Trung bình', 'Nặng', 'Rất nặng'],
    [['Trầm cảm', '24,4%', '12,41%', '8,47%', '1,17%', '1,17%'],
     ['Lo âu', '41,4%', '14,62%', '18,37%', '4,09%', '4,34%'],
     ['Căng thẳng', '15,6%', '8,58%', '4,20%', '2,14%', '0,66%']],
    widths=[2.5, 2.0, 2.5, 2.5, 2.5, 2.5])
doc.add_paragraph()

add_p('Các yếu tố nguy cơ: HS THPT (so với THCS), điểm dưới bách phân vị 60 (OR = 1,62), gia đình đơn thân, sống nội trú, sống ở nông thôn. Yếu tố bảo vệ mạnh: có người thân để tâm sự (OR = 0,22 cho trầm cảm, 0,27 cho lo âu, 0,23 cho căng thẳng). Giới tính không có ý nghĩa thống kê (p > 0,05).')

doc.add_paragraph()

# --- BÀI 10 ---
add_h('Bài 10: Salari và cs. (2024) — Tỷ lệ RLLAXA ở trẻ em, VTN và thanh niên: tổng quan hệ thống & phân tích tổng hợp', 2)
info_tbl([
    ('Tác giả', 'Nader Salari, Pegah Heidarian, Masoud Hassanabadi, Fateme Babajani, Nasrin Abdoli, Maliheh Aminian, Masoud Mohammadi'),
    ('Tạp chí', 'Journal of Prevention (IF ≈ 2.5)'),
    ('Năm', '2024'),
    ('Mẫu', '38 nghiên cứu (tổng quan hệ thống & phân tích tổng hợp)'),
    ('Công cụ', 'Random-effects meta-analysis'),
    ('Thiết kế', 'Tổng quan hệ thống và phân tích tổng hợp (Systematic Review & Meta-Analysis)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Tổng quan hệ thống và phân tích tổng hợp sử dụng mô hình hiệu ứng ngẫu nhiên (random-effects model) trên 38 nghiên cứu về rối loạn lo âu xã hội (Social Anxiety Disorder – SAD) ở trẻ em, vị thành niên và thanh niên trên toàn cầu.')

add_b('Kết quả chính')
add_p('Tỷ lệ rối loạn lo âu xã hội tăng dần theo lứa tuổi: trẻ em 4,7%, vị thành niên 8,3%, thanh niên 17,0%. Đây là bằng chứng mạnh cho thấy SAD gia tăng đáng kể qua các giai đoạn phát triển. SAD gây suy giảm nghiêm trọng tương tác xã hội, kết quả học tập và hoạt động nghề nghiệp.')
add_p('Các tác giả khuyến nghị các nhà hoạch định chính sách sử dụng kết quả này để xây dựng chiến lược phòng ngừa hiệu quả, đặc biệt nhắm vào nhóm vị thành niên và thanh niên là nhóm có nguy cơ cao nhất.')

doc.add_paragraph()

# --- BÀI 11 ---
add_h('Bài 11: Zhameden và cs. (2025) — Can thiệp tại trường học phòng ngừa lo âu và trầm cảm ở LMIC', 2)
info_tbl([
    ('Tác giả', 'Sharone Zhameden Dieu Yin, Mei Ken Low, Masuma Pervin Mishu'),
    ('Tạp chí', 'PLOS ONE'),
    ('Năm', '2025'),
    ('Mẫu', '6 RCTs, 1.587 người tham gia (tuổi 4–18)'),
    ('Công cụ', 'Tổng quan hệ thống RCTs'),
    ('Thiết kế', 'Tổng quan hệ thống (Systematic Review)')
])
doc.add_paragraph()

add_b('Phương pháp')
add_p('Tổng quan hệ thống các thử nghiệm ngẫu nhiên có đối chứng (RCTs) đánh giá hiệu quả can thiệp tại trường học nhằm phòng ngừa lo âu và trầm cảm ở trẻ em và vị thành niên tại các nước thu nhập thấp và trung bình (LMIC). Tổng cộng 6 RCTs với 1.587 người tham gia (4–18 tuổi).')

add_b('Kết quả chính')
add_p('Quốc gia nghiên cứu: Brazil, Trung Quốc, Liban, Malaysia (thu nhập trung bình cao) và Ấn Độ, Kenya (thu nhập trung bình thấp). Chưa có nghiên cứu nào từ nước thu nhập thấp. Hai loại can thiệp chính: CBT (4 nghiên cứu) và giáo dục tâm lý (2 nghiên cứu).')
add_p('Kết quả đáng chú ý: chỉ 1/4 nghiên cứu cho thấy giảm lo âu có ý nghĩa thống kê, trong khi 3/4 cho thấy cải thiện trầm cảm. Điều này gợi ý can thiệp tại trường có hiệu quả hơn cho trầm cảm so với lo âu. Chất lượng bằng chứng xếp hạng "rất thấp" (GRADE) do thiên vị phân bổ, thiếu mù hóa, và thiếu dữ liệu.')
add_p('Khuyến nghị: cần nghiên cứu ở nước thu nhập thấp, chuẩn hóa công cụ đo, kéo dài thời gian theo dõi trên 3 tháng, điều chỉnh can thiệp phù hợp văn hóa, và phát triển mô hình tiết kiệm chi phí.')

doc.add_paragraph()

# ==================== II. ĐIỂM NỔI BẬT ====================
add_h('II. ĐIỂM NỔI BẬT', 1)

add_b('1. Tỷ lệ lo âu ở Việt Nam thuộc nhóm cao nhất thế giới dù Đông Nam Á tỷ lệ thấp nhất')
add_p('Nghiên cứu GSHS trên 59 quốc gia cho thấy Đông Nam Á có tỷ lệ lo âu thấp nhất (3,78%). Tuy nhiên, các nghiên cứu riêng tại Việt Nam ghi nhận tỷ lệ 25,4–86,2%, cao hơn nhiều lần. Sự chênh lệch này có thể do khác biệt công cụ đo (GSHS vs DASS-21/GAD-7), ngưỡng chẩn đoán, và bối cảnh COVID-19. Điều này đặt ra câu hỏi về tính so sánh giữa các công cụ sàng lọc quốc tế và quốc gia.')

add_b('2. COVID-19 làm tăng đáng kể SKTT nhưng có xu hướng hồi phục')
add_p('Nghiên cứu trên 8.473 VTN Việt Nam cho thấy lo âu giảm từ 41,5% (trong dịch) xuống 25,4% (sau dịch), trầm cảm giảm từ 34,2% xuống 20,1%. Tuy nhiên, mức sau dịch vẫn cao, cho thấy hậu quả SKTT kéo dài. Đây là một trong số ít nghiên cứu so sánh trực tiếp hai giai đoạn tại Việt Nam với cỡ mẫu lớn.')

add_b('3. Học sinh dân tộc thiểu số – nhóm dễ bị tổn thương nhất')
add_p('Tỷ lệ lo âu 54,4% và trầm cảm 59,0% ở HS dân tộc thiểu số Lạng Sơn là mức cao đáng báo động. Gần một nửa (48,9%) từng trải qua ACEs. Nhóm này ít được nghiên cứu và ít tiếp cận dịch vụ SKTT, nhưng lại có nhu cầu can thiệp lớn nhất. Kết quả cho thấy chính sách SKTT học đường cần ưu tiên vùng dân tộc thiểu số.')

add_b('4. Căng thẳng học tập tăng dần theo năm học – bằng chứng từ nghiên cứu dọc')
add_p('Nghiên cứu dọc 3 năm tại Huế cho thấy điểm ESSA tăng từ 46,4 lên 53,5 (tăng 15,3%). Đây là bằng chứng hiếm hoi từ nghiên cứu dọc tại Việt Nam, xác nhận căng thẳng tích lũy theo thời gian. Yếu tố "học thêm" có tác động mạnh nhất (β = 4,73), gợi ý văn hóa học thêm tại Việt Nam có thể gây hại.')

add_b('5. Giới tính nữ là yếu tố nguy cơ nhất quán nhất')
add_p('Hầu hết các nghiên cứu xác nhận nữ giới có tỷ lệ lo âu, trầm cảm và căng thẳng cao hơn nam giới. Tại Việt Nam, Đông Nam Á và toàn cầu, xu hướng này nhất quán. Ngoại lệ duy nhất là nghiên cứu tại Trung Quốc (Dong và cs.) không tìm thấy sự khác biệt có ý nghĩa theo giới. Nguyên nhân có thể liên quan đến áp lực xã hội, kỳ vọng giới, và cơ chế sinh học.')

add_b('6. Can thiệp tại trường hiệu quả cho trầm cảm nhưng kém cho lo âu')
add_p('Tổng quan hệ thống 6 RCTs tại LMIC cho thấy 3/4 nghiên cứu cải thiện trầm cảm, nhưng chỉ 1/4 giảm lo âu. Điều này gợi ý lo âu có cơ chế phức tạp hơn, đòi hỏi can thiệp đa tầng (không chỉ CBT tại trường). Đặc biệt, chưa có RCT nào từ nước thu nhập thấp, cho thấy khoảng trống lớn về bằng chứng.')

add_b('7. Mối quan hệ gia đình là yếu tố bảo vệ mạnh nhất')
add_p('Tại Trung Quốc, HS có người thân tâm sự có nguy cơ trầm cảm thấp hơn 78% (OR = 0,22), lo âu thấp hơn 73% (OR = 0,27). Tại Việt Nam, chất lượng quan hệ cha mẹ–con là yếu tố bảo vệ hàng đầu. Phát hiện này nhất quán ở mọi bối cảnh văn hóa, gợi ý can thiệp gia đình nên là trụ cột trong mọi chương trình SKTT học đường.')

doc.add_paragraph()

# ==================== III. PHẢN BIỆN ====================
add_h('III. PHẢN BIỆN', 1)

add_b('Bài 1 – Phạm Thị Thu Hoa và cs. (2024)')
add_p('Ưu điểm: cỡ mẫu lớn (n = 3.910), phương pháp hỗn hợp, công cụ GAD-7 đã được chuẩn hóa. Hạn chế: nghiên cứu chỉ tại Hà Nội, thời điểm khảo sát (lockdown COVID-19) có thể phóng đại tỷ lệ lo âu. Mẫu thiên lệch vùng đô thị, không đại diện cho học sinh nông thôn. Phần định tính chỉ có 20 HS, chưa đủ để tổng quát hóa chiến lược ứng phó. Cần lặp lại nghiên cứu ở các tỉnh khác và so sánh trước-sau COVID.')

add_b('Bài 2 – Ngô Anh Vinh và cs. (2024)')
add_p('Ưu điểm: tập trung vào nhóm dân tộc thiểu số ít được nghiên cứu, kết hợp đánh giá ACEs. Hạn chế: chỉ có HS nội trú tại Lạng Sơn, không đại diện cho toàn bộ DTTS Việt Nam. DASS-21 chưa được chuẩn hóa đặc thù cho DTTS, có thể có lệch văn hóa. Thiếu nhóm đối chứng (HS Kinh cùng khu vực). Nghiên cứu cắt ngang không thể xác định quan hệ nhân quả ACEs→lo âu.')

add_b('Bài 3 – Trần Thảo Vi và cs. (2024)')
add_p('Ưu điểm hiếm có: thiết kế dọc 3 năm, theo dõi sự thay đổi theo thời gian. Hạn chế: mất mẫu đáng kể (611→341, mất 44,2%), có thể gây thiên lệch chọn. Chỉ ở Huế, 4 trường. Thang ESSA đo căng thẳng học tập chứ không phải lo âu lâm sàng. Lần cuối trùng COVID-19, khó tách tác động dịch với xu hướng tự nhiên. Tạp chí IF thấp (~1,0).')

add_b('Bài 4 – Nguyễn Ngọc Bảo Quyên và cs. (2025)')
add_p('Hạn chế lớn nhất: tỷ lệ lo âu 86,2% cao bất thường, có thể do khảo sát trực tuyến (tự chọn, thiên lệch chọn mẫu) và ngưỡng cắt DASS-21 quá thấp. Cỡ mẫu nhỏ (n = 501) và chỉ tại Hà Nội. Tạp chí trong nước, chưa qua peer-review quốc tế. Kết quả nên được diễn giải cẩn thận: DASS-21 sàng lọc chứ không chẩn đoán. Cần nghiên cứu xác nhận với phỏng vấn lâm sàng.')

add_b('Bài 5 – Hoàng Trung Học & Nguyễn Thùy Dung (2025)')
add_p('Ưu điểm: cỡ mẫu rất lớn (n = 8.473), 6 tỉnh, so sánh hai giai đoạn. Hạn chế: hai nhóm so sánh có thể là hai mẫu khác nhau (không phải theo dõi cùng nhóm), nên không thể kết luận "giảm" mà chỉ là "khác biệt giữa hai đợt khảo sát". Tạp chí quốc tế nhưng IF không rõ. Thiếu phân tích đa biến chi tiết cho từng tỉnh.')

add_b('Bài 6 – Sarfika và cs. (2025)')
add_p('Ưu điểm: đa trung tâm, áp dụng khung lý thuyết SCT. Hạn chế: chi tiết phương pháp và cỡ mẫu chưa rõ ràng trong bản tóm tắt. Chưa có thông tin về công cụ đo cụ thể và tỷ lệ đáp ứng. Kết quả giải thích phương sai cao (42,8–57,1%) nhưng cần kiểm tra overfitting. Cần truy cập toàn văn để đánh giá đầy đủ.')

add_b('Bài 7 – Ayeras và cs. (2024)')
add_p('Hạn chế rõ nhất: cỡ mẫu rất nhỏ (n = 12), chỉ tại một trường. Nghiên cứu hiện tượng học không thể tổng quát hóa. Tạp chí IF rất thấp. Tuy nhiên, cung cấp góc nhìn sâu về trải nghiệm chủ quan của HS có lo âu xã hội, bổ sung cho nghiên cứu định lượng. Cần nghiên cứu định lượng lớn tại Philippines để xác nhận.')

add_b('Bài 8 – Islam và cs. (2026)')
add_p('Ưu điểm vượt trội: cỡ mẫu khổng lồ (n = 179.937), 59 quốc gia, dữ liệu GSHS chuẩn hóa. Hạn chế: dữ liệu GSHS thu thập nhiều năm khác nhau, một số quốc gia có dữ liệu cũ. Tỷ lệ thấp ở Đông Nam Á (3,78%) mâu thuẫn với các nghiên cứu riêng lẻ tại Việt Nam, Indonesia, có thể do GSHS sử dụng câu hỏi sàng lọc đơn giản thay vì thang đo chuẩn hóa.')

add_b('Bài 9 – Dong, Wang & Lin (2025)')
add_p('Ưu điểm: cỡ mẫu tốt (n = 2.716), phân tích đa biến chi tiết. Hạn chế: chỉ tại Ya\'an (một thành phố), nghiên cứu cắt ngang. Thời điểm tháng 3–6/2022 trùng với chính sách Zero-COVID Trung Quốc, có thể phóng đại kết quả. Kết quả giới tính không có ý nghĩa – khác với hầu hết nghiên cứu khác, cần giải thích thêm.')

add_b('Bài 10 – Salari và cs. (2024)')
add_p('Ưu điểm: meta-analysis 38 nghiên cứu, phương pháp random-effects. Hạn chế: tập trung vào SAD (lo âu xã hội), không bao phủ các dạng lo âu khác (lo âu lan tỏa, lo âu thi cử). Thiếu thông tin về tính không đồng nhất (I²) trong bản tóm tắt. IF tạp chí trung bình. Cần bổ sung meta-analysis cho GAD, lo âu thi cử riêng biệt.')

add_b('Bài 11 – Zhameden và cs. (2025)')
add_p('Ưu điểm: tập trung LMIC, đánh giá GRADE. Hạn chế nghiêm trọng: chỉ 6 RCTs, chất lượng "rất thấp". Không có nghiên cứu từ nước thu nhập thấp. 4/6 dùng CBT – thiếu đa dạng phương pháp. Chưa có RCT nào từ Việt Nam hay Đông Nam Á ngoài Malaysia. Khoảng trống bằng chứng rất lớn, mở ra cơ hội nghiên cứu quan trọng.')

doc.add_paragraph()

# ==================== IV. HƯỚNG NGHIÊN CỨU ====================
add_h('IV. HƯỚNG NGHIÊN CỨU TƯƠNG LAI', 1)

add_b('A. Nhóm phương pháp luận')

add_p('1. Nghiên cứu dọc quy mô lớn tại Việt Nam: Thiếu vắng nghiên cứu dọc theo dõi HS từ THCS đến THPT (4–6 năm) với cỡ mẫu đại diện quốc gia. Nghiên cứu dọc cho phép xác định quan hệ nhân quả và điểm chuyển tiếp nguy cơ. Đề xuất: theo dõi ≥2.000 HS tại 10+ tỉnh, đo mỗi 6 tháng.')

add_p('2. So sánh công cụ đo lường: Các nghiên cứu sử dụng GAD-7, DASS-21, ESSA, GSHS cho kết quả rất khác nhau (3,78% vs 86,2%). Cần nghiên cứu so sánh trực tiếp các công cụ trên cùng một mẫu HS Việt Nam, xác định ngưỡng cắt phù hợp cho quần thể Việt.')

add_p('3. Nghiên cứu hỗn hợp sâu (mixed-methods): Kết hợp dữ liệu định lượng lớn với phỏng vấn sâu, nhật ký lo âu, và quan sát lớp học. Phương pháp này giúp hiểu "tại sao" và "như thế nào" HS lo âu, không chỉ "bao nhiêu phần trăm".')

add_b('B. Nhóm quần thể đặc biệt')

add_p('4. Học sinh dân tộc thiểu số – mở rộng: Chỉ có 1 nghiên cứu (Lạng Sơn). Cần mở rộng ra các tỉnh Tây Nguyên, Tây Bắc, Đồng bằng sông Cửu Long. So sánh DTTS vs Kinh cùng khu vực, kiểm soát SES. Đánh giá rào cản ngôn ngữ khi sử dụng công cụ sàng lọc tiếng Việt.')

add_p('5. Học sinh LGBT+ và lo âu: Chưa có nghiên cứu nào trong 11 bài đề cập HS LGBT+. Nghiên cứu quốc tế cho thấy nhóm này có tỷ lệ lo âu cao gấp 2–3 lần. Đây là khoảng trống lớn tại Việt Nam và Đông Nam Á.')

add_p('6. Học sinh khuyết tật và lo âu: Nhóm HS khuyết tật (học tập, thể chất, giác quan) hầu như vắng bóng trong nghiên cứu lo âu tại khu vực. Cần nghiên cứu đánh giá tác động kép: khuyết tật + lo âu.')

add_b('C. Nhóm yếu tố nguyên nhân')

add_p('7. Tác động của mạng xã hội và thiết bị điện tử: Nhiều nghiên cứu xác nhận liên quan giữa sử dụng thiết bị điện tử quá mức và lo âu, nhưng chưa có nghiên cứu can thiệp (RCT giảm thời gian sử dụng → đo lo âu). Đề xuất: RCT "digital detox" tại trường học Việt Nam.')

add_p('8. Áp lực học thêm – yếu tố đặc thù Việt Nam: Nghiên cứu tại Huế cho thấy "học thêm" có β cao nhất (4,73). Cần nghiên cứu sâu về cơ chế: học thêm làm tăng lo âu qua quá tải nhận thức, giảm thời gian nghỉ, hay tăng so sánh xã hội?')

add_p('9. Vai trò của ACEs ở quần thể chung: Nghiên cứu ACEs mới chỉ ở DTTS. Cần mở rộng đánh giá ACEs ở HS Kinh tại đô thị và nông thôn, đặc biệt các hình thức bạo lực gia đình tinh vi hơn (bạo lực tinh thần, kỳ vọng quá mức).')

add_b('D. Nhóm can thiệp')

add_p('10. RCT can thiệp tại trường học Việt Nam: Chưa có RCT nào về can thiệp SKTT tại trường học Việt Nam. Đề xuất: thử nghiệm CBT nhóm, mindfulness, hoặc giáo dục cảm xúc–xã hội (SEL) tại 10+ trường THCS/THPT, theo dõi 6–12 tháng.')

add_p('11. Can thiệp gia đình: Yếu tố gia đình là bảo vệ mạnh nhất (OR = 0,22). Cần nghiên cứu can thiệp: đào tạo kỹ năng giao tiếp cho phụ huynh, chương trình phối hợp gia đình–nhà trường, hỗ trợ phụ huynh nhận biết dấu hiệu lo âu.')

add_p('12. Can thiệp dựa trên công nghệ: Ứng dụng di động (app) tự hướng dẫn CBT, chatbot hỗ trợ SKTT, telepsychology cho HS vùng sâu vùng xa. Đặc biệt phù hợp với Việt Nam – tỷ lệ smartphone cao ngay cả ở nông thôn.')

add_p('13. Đánh giá hiệu quả kinh tế (cost-effectiveness): Chưa có nghiên cứu nào đánh giá chi phí–lợi ích của can thiệp SKTT học đường tại Việt Nam. Đây là thông tin cần thiết để thuyết phục chính sách đầu tư.')

doc.add_paragraph()

# ==================== V. DANH SÁCH BÀI CẦN ĐỌC ====================
add_h('V. DANH SÁCH BÀI CẦN ĐỌC THÊM', 1)

add_b('Nhóm A: Dịch tễ học lo âu ở VTN Đông Nam Á và LMIC')
tbl(['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['1', 'Anderson C.A.', '2025', 'J Child Adolesc Psychiatric Nursing', 'Tổng quan tường thuật về các yếu tố góp phần vào gia tăng lo âu ở VTN. Phân tích 61 bài báo, xác nhận tỷ lệ tăng so với trước đại dịch. Thế hệ Z có tỷ lệ lo âu cao nhất trong 3 thế hệ gần đây.'],
        ['2', 'Lancet Regional Health SEA', '2025', 'Lancet Reg Health SEA', 'Tổng quan hệ thống về bản chất, tỷ lệ và yếu tố quyết định SKTT ở VTN Nam Á. Tỷ lệ lo âu dao động 1,5–81,6% tùy phương pháp. Đánh giá khoảng trống nghiên cứu khu vực.'],
        ['3', 'UNICEF Việt Nam', '2022', 'UNICEF Report', 'Nghiên cứu về yếu tố trường học ảnh hưởng SKTT VTN Việt Nam. 1/5 VTN có vấn đề SKTT, chỉ 8,4% tiếp cận dịch vụ hỗ trợ. Chỉ 5,1% phụ huynh nhận ra con cần giúp đỡ.'],
    ],
    widths=[0.6, 2.5, 0.8, 2.8, 9.3])

doc.add_paragraph()
add_b('Nhóm B: Phương pháp can thiệp tại trường học')
tbl(['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['4', 'Cai, Mei, Wang & Luo', '2025', 'Frontiers Psychiatry (IF 4.7)', 'Meta-analysis 38 RCTs (15.730 người) về can thiệp tại trường tăng cường phục hồi cho trẻ/VTN. Hiệu quả nhỏ nhưng có ý nghĩa (SMD = 0,17). CBT, mindfulness, yoga đều có hiệu quả.'],
        ['5', 'Papageorgiou và cs.', '2025', 'Health Education Journal', 'Can thiệp giáo dục tâm lý tại trường phòng ngừa lo âu và căng thẳng dưới lâm sàng ở VTN. RCT tại Cyprus. CBT giảm lo âu ngắn hạn nhưng không duy trì sau 12 tháng.'],
        ['6', 'Frontiers Child Adolesc Psychiatry', '2025', 'Frontiers', 'Meta-analysis các can thiệp phổ quát tại trường cải thiện kết quả cảm xúc ở trẻ/thanh niên. So sánh targeted vs universal programs.'],
    ],
    widths=[0.6, 2.5, 0.8, 2.8, 9.3])

doc.add_paragraph()
add_b('Nhóm C: Yếu tố nguy cơ và bảo vệ')
tbl(['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['7', 'Nguyễn Thị Thùy', '2023', 'Open Psychology J', 'Đặc điểm nhân cách và RLLA ở VTN Việt Nam: vai trò trung gian của hỗ trợ xã hội và lòng tự trọng. n=582 HS THCS. Loạn thần kinh tương quan r=0,56 với lo âu.'],
        ['8', 'Giang Thiên Vũ và cs.', '2023', 'HCMUE J Science (Vol. 20, No. 8)', 'SKTT VTN TP.HCM: sàng lọc nguy cơ và giải pháp. n=6.575, 21,13% lo âu nặng/rất nặng. Áp lực thi cử, thiếu hướng nghiệp là yếu tố chính.'],
        ['9', 'Nghiên cứu ResearchGate', '2024', 'ResearchGate', 'Áp lực học tập và lo âu ở HS THPT TP. Thủ Đức, TP.HCM. n=421, lo âu 35,6%. Mối liên quan áp lực→lo âu có ý nghĩa (p<0,001).'],
    ],
    widths=[0.6, 2.5, 0.8, 2.8, 9.3])

doc.add_paragraph()
add_b('Nhóm D: Xu hướng toàn cầu và so sánh liên quốc gia')
tbl(['#', 'Tác giả', 'Năm', 'Tạp chí', 'Nội dung'],
    [
        ['10', 'ScienceDirect', '2025', 'J Affective Disorders', 'Xu hướng triệu chứng trầm cảm và lo âu ở VTN 10–24 tuổi từ 1990–2021: phân tích GBD. Tỷ lệ tăng đều qua 3 thập kỷ, đặc biệt nhanh ở nữ và nhóm 15–19 tuổi.'],
        ['11', 'Prevalence anxiety/depression LMIC', '2025', 'Psychiatric Research & Clinical Practice', 'Tỷ lệ lo âu và trầm cảm ở trẻ/VTN tại LMIC: tổng quan hệ thống. Tỷ lệ 1–30% tùy khu vực. Nhấn mạnh khoảng trống bằng chứng tại châu Phi và Đông Nam Á.'],
    ],
    widths=[0.6, 2.5, 0.8, 2.8, 9.3])

doc.add_paragraph()

# ==================== BỔ SUNG 10 BÀI MỚI ====================
doc.add_page_break()
add_h('PHẦN D: 10 BÀI NGHIÊN CỨU BỔ SUNG (5 VN + 3 ĐNA + 2 THẾ GIỚI)', 1)

# === VIỆT NAM BỔ SUNG ===
add_h('D.1 VIỆT NAM BỔ SUNG', 2)

# --- BÀI 12 ---
add_h('Bài 12: Sàng lọc lo âu, trầm cảm, stress ở HS THPT Long Bình, An Giang (2024)', 3)
info_tbl([
    ('Tác giả', 'Nhóm nghiên cứu An Giang'),
    ('Tạp chí', 'Tạp chí Y học Việt Nam (2025)'),
    ('Năm', '2025 (khảo sát tháng 6/2024)'),
    ('Mẫu', '366 học sinh THPT Long Bình, An Giang'),
    ('Công cụ', 'DASS-21'),
    ('Thiết kế', 'Nghiên cứu cắt ngang')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Lo âu: 61,2% (nhẹ 9,3%, vừa 24,0%, nặng 12,6%, rất nặng 15,3%). Trầm cảm: 47,3% (nhẹ 15,8%, vừa 18,0%, nặng 8,5%, rất nặng 4,9%). Stress: 38,0% (nhẹ 12,8%, vừa 16,4%, nặng 11,5%, rất nặng 4,9%).')
add_p('Phân bố đáng chú ý: 28,1% không có triệu chứng; 24,6% mắc 1 vấn đề; 20,0% mắc 2 vấn đề; 27,3% mắc cả 3 vấn đề cùng lúc. Tỷ lệ rất cao so với các nghiên cứu ở đô thị, có thể phản ánh đặc thù vùng nông thôn và thiếu tiếp cận dịch vụ SKTT.')

doc.add_paragraph()

# --- BÀI 13 ---
add_h('Bài 13: Trần Hồ Vĩnh Lộc và cs. (2024) — Trầm cảm, lo âu, stress ở HS THPT tại TP.HCM', 3)
info_tbl([
    ('Tác giả', 'Trần Hồ Vĩnh Lộc, Huỳnh Ngọc Vân Anh, Tô Gia Kiên'),
    ('Tạp chí', 'Tạp chí Y học TP.HCM (Vol. 27, No. 5)'),
    ('Năm', '2024 (khảo sát tháng 2–4/2024)'),
    ('Mẫu', '976 học sinh THPT tại 2 trường ở TP.HCM (tuổi TB 17; 53,4% nữ)'),
    ('Công cụ', 'DASS-Y (dành cho thanh thiếu niên); ESSA (áp lực học tập)'),
    ('Thiết kế', 'Nghiên cứu cắt ngang mô tả, chọn mẫu ngẫu nhiên nhiều giai đoạn')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Trầm cảm 31,7%, lo âu 25,1%, stress 23,8%. Tổng cộng 42,4% HS mắc ít nhất 1 vấn đề SKTT, 13,2% mắc cả 3. Nữ giới có tỷ lệ cao hơn nam ở cả 3 chỉ số. Áp lực học tập (vừa hoặc nặng) là yếu tố dự báo mạnh nhất.')
add_p('Cha mẹ ly hôn/ly thân và tình trạng kinh tế gia đình thấp làm tăng nguy cơ trầm cảm và stress đáng kể. Thú vị: trình độ học vấn mẹ cao hơn tương quan với tỷ lệ stress cao hơn — có thể do kỳ vọng cao hơn.')

doc.add_paragraph()

# --- BÀI 14 ---
add_h('Bài 14: Nguyễn Thị Hương và cs. (2024) — Trầm cảm ở HS THCS Chu Văn An, Thái Nguyên', 3)
info_tbl([
    ('Tác giả', 'Nguyễn Thị Hương, Trần Nguyên Ngọc, Lê Thị Thủy Linh'),
    ('Tạp chí', 'Tạp chí Y học Việt Nam (Vol. 543, No. 2)'),
    ('Năm', '2024 (năm học 2023–2024)'),
    ('Mẫu', '492 học sinh THCS Chu Văn An, Thái Nguyên (tuổi TB 13,3; 55,7% nữ)'),
    ('Công cụ', 'Chẩn đoán lâm sàng theo ICD-10'),
    ('Thiết kế', 'Nghiên cứu cắt ngang mô tả')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Tỷ lệ trầm cảm theo chẩn đoán ICD-10: 6,3% (31 HS). Đây là tỷ lệ chẩn đoán lâm sàng, thấp hơn nhiều so với tỷ lệ sàng lọc DASS-21 (thường 30–80%), cho thấy tầm quan trọng của việc phân biệt sàng lọc và chẩn đoán.')
add_p('Trong số HS trầm cảm: nhẹ 51,16%, trung bình 22,6%, rối loạn thích ứng có phản ứng trầm cảm 22,6%. Triệu chứng thường gặp: rối loạn giấc ngủ (29,1%), khó tập trung (11,2%), ý tưởng tự gây hại (4,7%). Nữ chiếm đa số.')

doc.add_paragraph()

# --- BÀI 15 ---
add_h('Bài 15: Nguyễn HTT và cs. (2024) — Trầm cảm ở HS THPT đô thị Hà Nội: vai trò quan hệ gia đình–trường học', 3)
info_tbl([
    ('Tác giả', 'Nguyễn HTT, Trần BX, Lưu HN, Boyer L, Fond G, Auquier P, Latkin CA, Nguyễn TT, Zhang MWB, Ho RCM, Ho CSH'),
    ('Tạp chí', 'Journal of Epidemiology and Population Health'),
    ('Năm', '2024'),
    ('Mẫu', '507 học sinh 15–17 tuổi tại Hà Nội'),
    ('Công cụ', 'RADS-2 (Reynolds Adolescent Depression Scale, 2nd Edition)'),
    ('Thiết kế', 'Nghiên cứu cắt ngang, phân tích hồi quy Tobit đa biến')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Đánh giá trầm cảm qua 4 miền RADS-2: tâm trạng buồn bã (15,1 ± 4,2), mất hứng thú (16,4 ± 4,0), tự đánh giá tiêu cực (13,1 ± 4,4), than phiền cơ thể (12,4 ± 3,7). Không sử dụng ngưỡng cắt mà phân tích liên tục — cách tiếp cận tinh tế hơn.')
add_p('Yếu tố nguy cơ: quan hệ cha mẹ kém, thiếu người tâm sự, xung đột gia đình, áp lực thi cử, quá tải bài vở. Yếu tố bảo vệ: hài lòng với bạn bè ở trường, hỗ trợ từ giáo viên/bạn bè, tham gia hoạt động ngoại khóa.')

doc.add_paragraph()

# --- BÀI 16 ---
add_h('Bài 16: Nguyễn Danh Lâm và cs. (2024) — Stress, lo âu, trầm cảm ở HS THPT Yên Định, Thanh Hóa', 3)
info_tbl([
    ('Tác giả', 'Nguyễn Danh Lâm, Lê Minh Giang, Nguyễn Thị Phương Mai, Nguyễn Thị Diệu Thúy, Nguyễn Thị Thanh Mai'),
    ('Tạp chí', 'Tạp chí Y học Việt Nam (Vol. 516, No. 1)'),
    ('Năm', '2024'),
    ('Mẫu', '482 học sinh THPT huyện Yên Định, Thanh Hóa'),
    ('Công cụ', 'DASS-21'),
    ('Thiết kế', 'Nghiên cứu cắt ngang mô tả')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Stress 41,7%, lo âu 49,0%, trầm cảm 43,6% — chủ yếu ở mức nhẹ và trung bình. Tỷ lệ đáng lo ngại: khoảng 1/3 HS có ý tưởng tự gây hại, 10% tự gây thương tích, 25% từng nghĩ đến tự tử, 1,4% đã cố gắng tự tử.')
add_p('Đây là nghiên cứu hiếm hoi tại vùng bán đô thị (huyện Yên Định) — khu vực ít được nghiên cứu so với thành phố lớn. Tác giả nhấn mạnh nhu cầu cấp thiết về hệ thống tư vấn SKTT tại trường học khu vực nông thôn và bán đô thị.')

doc.add_paragraph()

# === ĐÔNG NAM Á BỔ SUNG ===
add_h('D.2 ĐÔNG NAM Á BỔ SUNG', 2)

# --- BÀI 17 ---
add_h('Bài 17: GBD 2021 — Dịch tễ học và gánh nặng 10 rối loạn tâm thần tại ASEAN (2025)', 3)
info_tbl([
    ('Tác giả', 'Nhóm nghiên cứu GBD 2021 (đa quốc gia)'),
    ('Tạp chí', 'The Lancet Public Health (IF ≈ 72.4)'),
    ('Năm', '2025'),
    ('Mẫu', '10 quốc gia ASEAN (Brunei, Cambodia, Indonesia, Lào, Malaysia, Myanmar, Philippines, Singapore, Thái Lan, Việt Nam), dữ liệu GBD 1990–2021'),
    ('Công cụ', 'Mô hình GBD (Global Burden of Disease)'),
    ('Thiết kế', 'Phân tích dữ liệu thứ cấp GBD 2021')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Tỷ lệ rối loạn tâm thần chuẩn hóa theo tuổi tại ASEAN: 11,9% (10,9–12,9) năm 2021. Lo âu và trầm cảm là phổ biến nhất. Dao động từ 10,1% (Việt Nam) đến 13,2% (Malaysia). Tổng số ca: 80,4 triệu (tăng 70% so với 1990).')
add_p('DALYs: 11,2 triệu (tăng 87,4% từ 1990). Nhóm 10–14 tuổi có gánh nặng bệnh tật cao nhất — rối loạn tâm thần chiếm 16,3% tổng DALYs ở nhóm tuổi này. Đây là bằng chứng cấp Lancet cho thấy vị thành niên ASEAN chịu gánh nặng SKTT nặng nề nhất.')

doc.add_paragraph()

# --- BÀI 18 ---
add_h('Bài 18: So sánh giáo dục SKTT tại trường học ở Philippines, Indonesia và Nhật Bản (2025)', 3)
info_tbl([
    ('Tác giả', 'Nhóm nghiên cứu đa quốc gia (Philippines, Indonesia, Nhật Bản)'),
    ('Tạp chí', 'Tropical Medicine and Health'),
    ('Năm', '2025'),
    ('Mẫu', 'Phân tích chính sách & chương trình giảng dạy: Philippines 22 tài liệu, Indonesia 9, Nhật Bản 6'),
    ('Công cụ', 'Phân tích so sánh chính sách và chương trình giảng dạy'),
    ('Thiết kế', 'Nghiên cứu so sánh (comparative study)')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Philippines tích hợp SKTT vào giáo dục sức khỏe, giáo dục giá trị và hướng dẫn lớp chủ nhiệm, bao phủ gần như toàn bộ lớp 1–12. Nhật Bản tích hợp vào giáo dục sức khỏe và thể chất lớp 5–10. Indonesia tích hợp vào giáo dục tôn giáo lớp 1–12.')
add_p('Philippines và Indonesia có chính sách SKTT rõ ràng, nhưng Nhật Bản không có chính sách lõi về giáo dục SKTT mà phát triển qua hướng dẫn chương trình học. Cả 3 nước đều thực hiện giáo dục SKTT chưa đầy đủ. Phát hiện quan trọng: Việt Nam và các nước ASEAN khác vắng bóng trong nghiên cứu so sánh này.')

doc.add_paragraph()

# --- BÀI 19 ---
add_h('Bài 19: Mallari và cs. (2025) — Chương trình SKTT trường học và yếu tố bảo vệ ở VTN Philippines', 3)
info_tbl([
    ('Tác giả', 'Mallari và cộng sự'),
    ('Tạp chí', 'Psychology in the Schools (Wiley)'),
    ('Năm', '2025'),
    ('Mẫu', 'Vị thành niên đi học tại Philippines'),
    ('Công cụ', 'Đánh giá chương trình SKTT trường học + yếu tố bảo vệ'),
    ('Thiết kế', 'Nghiên cứu cắt ngang')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Tỷ lệ rối nhiễu tâm lý gia tăng ở Philippines — quốc gia có nguồn lực SKTT rất hạn chế. 33% VTN Philippines báo cáo rối nhiễu tâm lý nhưng chỉ 2% sử dụng dịch vụ hỗ trợ SKTT. Khoảng cách dịch vụ rất lớn.')
add_p('Các yếu tố bảo vệ được xác định: hỗ trợ xã hội, khả năng phục hồi, hoạt động thể chất, và các chương trình SKTT tại trường. Nghiên cứu nhấn mạnh sự cần thiết phải mở rộng chương trình SKTT trường học với nguồn lực hạn chế — bài học cho Việt Nam và các nước ASEAN tương tự.')

doc.add_paragraph()

# === THẾ GIỚI BỔ SUNG ===
add_h('D.3 THẾ GIỚI BỔ SUNG', 2)

# --- BÀI 20 ---
add_h('Bài 20: Meta-analysis can thiệp toàn trường (whole-school) nâng cao SKTT ở VTN (2025)', 3)
info_tbl([
    ('Tác giả', 'Nhóm nghiên cứu quốc tế'),
    ('Tạp chí', 'Journal of Youth and Adolescence (Springer, IF ≈ 4.6)'),
    ('Năm', '2025'),
    ('Mẫu', '28 nghiên cứu trong 58 ấn phẩm'),
    ('Công cụ', 'Tổng quan hệ thống và phân tích tổng hợp (PRISMA)'),
    ('Thiết kế', 'Tổng quan hệ thống và phân tích tổng hợp RCTs cập nhật')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('Đánh giá hiệu quả can thiệp "whole-school" (toàn trường) — cách tiếp cận tích hợp SKTT vào mọi hoạt động trường học thay vì chỉ can thiệp cho cá nhân/nhóm nhỏ. 28 nghiên cứu gồm các chương trình đa dạng: SEL (học tập cảm xúc–xã hội), phòng chống bắt nạt, tăng cường phục hồi.')
add_p('Kết quả cho thấy can thiệp whole-school có hiệu quả cho cả HS khỏe mạnh và HS có triệu chứng SKTT cao. CBT là phương thức hiệu quả nhất cho giảm lo âu. Tuy nhiên, hiệu quả duy trì dài hạn (>12 tháng) vẫn chưa rõ ràng — cần booster sessions. Bài học cho Việt Nam: nên phát triển mô hình whole-school thay vì chỉ tư vấn cá nhân.')

doc.add_paragraph()

# --- BÀI 21 ---
add_h('Bài 21: Anderson (2025) — Yếu tố góp phần vào gia tăng lo âu ở VTN: tổng quan tường thuật', 3)
info_tbl([
    ('Tác giả', 'C.A. Anderson'),
    ('Tạp chí', 'Journal of Child and Adolescent Psychiatric Nursing (Wiley)'),
    ('Năm', '2025'),
    ('Mẫu', 'Tổng quan 61 bài báo'),
    ('Công cụ', 'Narrative review'),
    ('Thiết kế', 'Tổng quan tường thuật (Narrative Review)')
])
doc.add_paragraph()

add_b('Kết quả chính')
add_p('31,9% VTN 13–18 tuổi đã hoặc đang được chẩn đoán rối loạn lo âu. 31,5% có triệu chứng trầm cảm. Thế hệ Z (sinh 1997–2012) có tỷ lệ lo âu cao nhất trong 3 thế hệ gần đây. COVID-19 làm trầm trọng thêm xu hướng vốn đã tăng trước đại dịch.')
add_p('48/52 nghiên cứu cho thấy tương quan dương giữa áp lực học tập và SKTT kém ở VTN. Các yếu tố khác: mạng xã hội, thiếu ngủ, cô lập xã hội, thay đổi cấu trúc gia đình. Khuyến nghị: giáo dục phụ huynh, chương trình SEL tại trường, tăng cường kỹ năng ứng phó cho VTN.')

doc.add_paragraph()

# ==================== CẬP NHẬT PHẢN BIỆN CHO 10 BÀI MỚI ====================
add_h('PHẢN BIỆN 10 BÀI BỔ SUNG', 1)

add_b('Bài 12 – An Giang (2024)')
add_p('Hạn chế: cỡ mẫu nhỏ (n = 366), chỉ 1 trường, vùng nông thôn An Giang không đại diện cả nước. Tỷ lệ lo âu 61,2% cao bất thường — cần xem xét ngưỡng cắt DASS-21 và điều kiện khảo sát. Ưu điểm: dữ liệu hiếm từ vùng nông thôn ĐBSCL.')

add_b('Bài 13 – Trần Hồ Vĩnh Lộc (TP.HCM, 2024)')
add_p('Ưu điểm: cỡ mẫu tốt (n = 976), chọn mẫu ngẫu nhiên nhiều giai đoạn, dùng DASS-Y (chuyên cho TTN). Hạn chế: chỉ 2 trường TP.HCM. Phát hiện thú vị về "trình độ mẹ cao → stress cao" cần nghiên cứu sâu hơn. Tạp chí trong nước.')

add_b('Bài 14 – Nguyễn Thị Hương (Thái Nguyên, 2024)')
add_p('Ưu điểm nổi bật: dùng chẩn đoán lâm sàng ICD-10 (không chỉ sàng lọc), cho tỷ lệ 6,3% — thực tế hơn nhiều so với DASS-21. Hạn chế: chỉ 1 trường THCS ở đô thị. Không đánh giá lo âu, chỉ trầm cảm. Bài tham chiếu quan trọng để so sánh sàng lọc vs chẩn đoán.')

add_b('Bài 15 – Nguyễn HTT (Hà Nội, 2024)')
add_p('Ưu điểm: hợp tác quốc tế (Singapore, Úc, Pháp), dùng RADS-2 chuyên biệt, phân tích Tobit đa biến. Hạn chế: chỉ HS Hà Nội, thiết kế cắt ngang. RADS-2 đo trầm cảm, không đo lo âu trực tiếp. Tuy nhiên, kết quả về vai trò gia đình/trường học rất giá trị cho thiết kế can thiệp.')

add_b('Bài 16 – Nguyễn Danh Lâm (Thanh Hóa, 2024)')
add_p('Ưu điểm: dữ liệu hiếm từ vùng bán đô thị. Dữ liệu về tự tử rất đáng lo. Hạn chế: bài đăng năm 2022 (Vol. 516) nên dữ liệu có thể từ 2021–2022, sát giới hạn 2024. Cỡ mẫu vừa (n = 482). DASS-21 là sàng lọc, tỷ lệ cao (49% lo âu) cần xác nhận bằng chẩn đoán.')

add_b('Bài 17 – GBD 2021 ASEAN (Lancet, 2025)')
add_p('Ưu điểm vượt trội: Lancet Public Health (IF ~72), dữ liệu GBD chuẩn hóa 30 năm, 10 quốc gia ASEAN. Hạn chế: dữ liệu GBD phụ thuộc vào chất lượng báo cáo quốc gia — Việt Nam, Myanmar, Lào có thể underreport. Không phân tích riêng HS THCS/THPT. Tuy nhiên, nhóm 10–14 tuổi có gánh nặng cao nhất = rất phù hợp chủ đề.')

add_b('Bài 18 – So sánh giáo dục SKTT 3 nước châu Á (2025)')
add_p('Ưu điểm: so sánh chính sách thực tế, ứng dụng cao. Hạn chế: chỉ 3 nước (thiếu VN, Thái Lan, Malaysia), phân tích tài liệu chứ không khảo sát thực tế. Không đánh giá hiệu quả can thiệp. Tuy nhiên, bổ sung góc nhìn chính sách mà các nghiên cứu dịch tễ thiếu.')

add_b('Bài 19 – Mallari (Philippines, 2025)')
add_p('Hạn chế: chi tiết phương pháp và cỡ mẫu chưa rõ từ bản tóm tắt. Tạp chí Psychology in the Schools có IF trung bình. Ưu điểm: dữ liệu về khoảng cách dịch vụ (33% có vấn đề nhưng chỉ 2% dùng dịch vụ) rất quan trọng cho chính sách ASEAN.')

add_b('Bài 20 – Whole-school meta-analysis (2025)')
add_p('Ưu điểm: meta-analysis cập nhật, 28 RCTs, so sánh targeted vs universal. Hạn chế: phần lớn nghiên cứu từ phương Tây (Mỹ, Úc, Anh), chưa rõ áp dụng được cho bối cảnh ASEAN. Cần RCT whole-school tại Việt Nam để xác nhận.')

add_b('Bài 21 – Anderson narrative review (2025)')
add_p('Ưu điểm: tổng hợp 61 bài, cho bức tranh toàn cảnh về gia tăng lo âu VTN toàn cầu. Hạn chế: tổng quan tường thuật (không phải hệ thống), không có phân tích tổng hợp định lượng. Dữ liệu chủ yếu từ phương Tây. Tuy nhiên, phát hiện "48/52 NC xác nhận áp lực học tập→SKTT kém" rất mạnh.')

doc.add_paragraph()

# ==================== FOOTER ====================
add_i('Báo cáo được tổng hợp ngày 27/03/2026 từ các nguồn: PubMed, PMC, Frontiers, PLOS ONE, Elsevier, Springer, Lancet, Wiley, UNICEF, tạp chí Y học Việt Nam.')
add_i('Tổng cộng: 21 bài (10 Việt Nam + 5 Đông Nam Á + 6 thế giới). Tất cả tỷ lệ là sàng lọc trừ Bài 14 (chẩn đoán ICD-10).')

# ==================== SAVE ====================
output_path = r'c:\Users\HLC\OneDrive\read_books\Lo-au\Lo âu - THCS THPT - 27032026.docx'
doc.save(output_path)
print('Done: ' + output_path)
