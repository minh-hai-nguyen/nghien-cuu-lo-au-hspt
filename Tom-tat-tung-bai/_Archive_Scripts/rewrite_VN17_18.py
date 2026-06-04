# -*- coding: utf-8 -*-
"""VN17 (Danh Lâm Thanh Hóa) + VN18 (An Giang) — chuẩn CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc
def rb(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(doc, t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(c, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); c._tc.get_or_add_tcPr().append(s)
def set_w(c, w):
    tcW = c._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ============================================================
# VN17: NGUYỄN DANH LÂM 2022 — Yên Định, Thanh Hóa
# ============================================================
d = make_doc()
bl(d, 'Tóm tắt bài VN-17', bold=True)

rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d, 'Công trình nghiên cứu \u00ab Thực trạng nguy cơ stress, lo âu, trầm cảm của học sinh trung học phổ thông huyện Yên Định, Thanh Hóa \u00bb (Risk of Stress, Anxiety, Depression of High School Students in Yen Dinh, Thanh Hoa), do Nguyễn Danh Lâm (Bệnh viện huyện Yên Định), Lê Minh Giang (Đại học Y Hà Nội), Nguyễn Thị Phương Mai, Nguyễn Thị Diệu Thúy và Nguyễn Thị Thanh Mai (2022), khảo sát trên mẫu 482 học sinh THPT tại 2 trường (THPT Yên Định I: 264 HS; THPT Thống Nhất: 218 HS), huyện Yên Định, tỉnh Thanh Hóa, thu thập tháng 9/2021, xuất bản trên Tạp chí Y học Việt Nam, tập 516, số 1, tháng 7/2022, trang 67\u201370.')

rb(d, 'Phương pháp nghiên cứu')
bl(d, 'Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21 \u2014 Depression Anxiety Stress Scale) phiên bản tiếng Việt. Nói cách khác, nghiên cứu sử dụng phương pháp sàng lọc cắt ngang bằng bảng câu hỏi tự điền tại trường học.', bold=True)

bl(d, 'DASS-21 là phiên bản rút gọn 21 mục của thang gốc 42 mục (Lovibond & Lovibond, 1995), đã được xác thực cho cộng đồng Việt Nam (Thạch và cộng sự, 2013). Thang điểm 4 mức (0\u20133), phân loại: không rối loạn, nhẹ, vừa, nặng, rất nặng.')

bl(d, 'Tổng quan tài liệu của chúng tôi cho thấy vùng bán đô thị Thanh Hóa \u2014 nơi có đặc điểm trung gian giữa nông thôn và thành thị \u2014 hầu như chưa được khảo sát về sức khỏe tâm thần thanh thiếu niên. Tỷ lệ tham chiếu: 41,1% trầm cảm ở THPT theo Nguyễn Tấn Đạt (2013). Đây là nghiên cứu đầu tiên kết hợp đánh giá DAS và hành vi tự hại tại khu vực này.', bold=True)

bl(d, 'Dữ liệu gồm thông tin cá nhân (bảng hỏi cấu trúc), tình trạng stress, lo âu, trầm cảm (DASS-21), và các câu hỏi về ý định tự làm đau bản thân và tự tử.')
bl(d, 'Cỡ mẫu 482 HS, tính theo công thức ước lượng 1 tỷ lệ (p = 0,411, \u03b5 = 0,04, \u03b1 = 0,05). Chọn mẫu có chủ đích, 2 trường THPT, HS tự điền bảng hỏi.')

rb(d, 'Kết quả nghiên cứu định lượng')
bl(d, 'Tỷ lệ stress, lo âu, trầm cảm (N = 482, DASS-21):', bold=True)
tbl(d,
    ['Tình trạng', 'Tỷ lệ tổng', 'Nhẹ', 'Vừa', 'Nặng', 'Rất nặng'],
    [['Stress', '41,7%', '16,0%', '14,1%', '6,8%', '4,8%'],
     ['Lo âu', '49,0%', '7,7%', '24,5%', '8,1%', '4,6%'],
     ['Trầm cảm', '43,6%', '12,2%', '27,2%', '2,9%', '1,2%']],
    widths=[2.5, 2.0, 1.5, 1.5, 1.5, 2.0])
d.add_paragraph()

bl(d, 'Hành vi tự hại và tự tử (N = 482):', bold=True)
tbl(d,
    ['Biểu hiện', 'Chưa bao giờ', 'Đã nghĩ đến', 'Đã thực hiện'],
    [['Tự làm đau bản thân', '292 (58,4%)', '158 (31,6%)', '50 (10,0%)'],
     ['Tự tử', '375 (75,0%)', '118 (23,6%)', '7 (1,4%)']],
    widths=[4.0, 3.0, 3.0, 3.0])
d.add_paragraph()

bh3(d, 'Phát hiện đặc biệt về tự hại : Có tới 1/3 số trẻ (31,6%) đã từng nghĩ đến tự làm đau bản thân và 10% đã thực sự tự làm đau. Nổi bật hơn, 1/4 số trẻ (23,6%) đã từng nghĩ đến tự tử và 1,4% đã cố tự tử nhưng không thành công. Đây là dữ liệu đáng báo động \u2014 ít nghiên cứu tại Việt Nam đánh giá hành vi tự hại trực tiếp.')

rb(d, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3(d, 'Phát hiện về tỷ lệ và hành vi tự hại ở vùng bán đô thị')
bl(d, '*Tỷ lệ ở mức trung bình-cao.* Lo âu 49,0% cao hơn Hoa 2024 (GAD-7: 40,6%) nhưng thấp hơn An Giang 2025 (DASS-21: 61,2%) và Bảo Quyên 2025 (DASS-21: 86,2%). Nói cách khác, tỷ lệ tại vùng bán đô thị Thanh Hóa nằm ở mức trung gian \u2014 phù hợp với đặc điểm kinh tế xã hội vùng.')
bl(d, '*Hành vi tự hại \u2014 dữ liệu quan trọng nhất.* 10% đã tự làm đau và 1,4% đã cố tự tử \u2014 đây là thông tin cấp thiết cho can thiệp. So sánh: UNICEF VN báo cáo 5,8% VTN đã cố tự tử. Dữ liệu này cho thấy vùng bán đô thị cũng cần quan tâm, không chỉ đô thị lớn.')
bl(d, '*Chủ yếu mức nhẹ-vừa.* Lo âu nặng + rất nặng chỉ 12,7%, trầm cảm nặng + rất nặng 4,1% \u2014 gợi ý can thiệp sớm ở giai đoạn nhẹ-vừa có thể hiệu quả cao.')

rb(d, 'Kết luận')
bl(d, 'Dữ liệu của chúng tôi, cho thấy tỷ lệ đáng kể stress (41,7%), lo âu (49,0%) và trầm cảm (43,6%) ở 482 HS THPT vùng bán đô thị Thanh Hóa, kèm theo 10% đã tự làm đau và 1,4% đã cố tự tử, gợi ý rằng sức khỏe tâm thần thanh thiếu niên không chỉ là vấn đề của đô thị lớn. Nhu cầu thiết lập hệ thống hướng dẫn kỹ năng ứng phó stress và tiếp cận dịch vụ tư vấn tại trường THPT vùng bán đô thị là cấp thiết.', bold=True)

rh2(d, 'Phản biện')
bl(d, 'Điểm mạnh: vùng bán đô thị Thanh Hóa ít NC, dữ liệu tự hại/tự tử quan trọng, DASS-21 đã xác thực tại VN. Tuy nhiên, chọn mẫu có chủ đích (2 trường) \u2014 không ngẫu nhiên, thiên lệch. Chỉ thống kê mô tả, KHÔNG có phân tích yếu tố liên quan (hồi quy) \u2014 không biết yếu tố nào ảnh hưởng. Bài chỉ 4 trang \u2014 thiếu thảo luận sâu, không phân tích giới tính, khối lớp. Tạp chí Y học VN không lập chỉ mục quốc tế.')

rh2(d, 'Hướng nghiên cứu tiếp theo')
bl(d, 'Phân tích yếu tố nguy cơ bằng hồi quy đa biến. So sánh giới tính và khối lớp. Nghiên cứu can thiệp giảm hành vi tự hại \u2014 ưu tiên cao. Mở rộng nhiều trường Thanh Hóa. Xuất bản trên tạp chí có chỉ mục QT.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50 Trung bình-Thấp. Vùng bán đô thị ít NC, dữ liệu tự hại quan trọng, nhưng chọn mẫu có chủ đích, chỉ mô tả, bài 4 trang, tạp chí không QT.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d.save('VN17_DanhLam_2022.docx')
sys.stderr.write('VN17 OK\n')

# ============================================================
# VN18: AN GIANG 2025 — Long Bình, ĐBSCL
# ============================================================
d2 = make_doc()
bl(d2, 'Tóm tắt bài VN-18', bold=True)

rb(d2, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d2, 'Công trình nghiên cứu \u00ab Kết quả sàng lọc lo âu, trầm cảm, stress bằng DASS-21 ở học sinh phổ thông Long Bình, tỉnh An Giang năm 2024 \u00bb (Results of Screening for Anxiety, Depression, and Stress Using DASS-21 Among Long Binh High School Students, An Giang Province in 2024), do Lê Minh T., Nguyễn Đăng K. và Ngô Anh V. (2025), khảo sát trên 366 học sinh THPT Long Bình, tỉnh An Giang, thu thập tháng 6/2024, xuất bản trên Tạp chí Y học Việt Nam, tập 549, số 1, tháng 4/2025, trang 32\u201335. DOI: 10.51298/vmj.v549i1.13506.')

rb(d2, 'Phương pháp nghiên cứu')
bl(d2, 'Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21 \u2014 Depression Anxiety Stress Scale) phiên bản tiếng Việt. Nói cách khác, nghiên cứu sử dụng phương pháp sàng lọc cắt ngang bằng bảng câu hỏi tại trường học.', bold=True)

bl(d2, 'DASS-21 phiên bản tiếng Việt đã được xác thực (Lê Thị Hồng Minh và Thạch Đức Trần, 2017; PLoS ONE). Phân tích bằng kiểm định chi bình phương. Nghiên cứu được phê duyệt đạo đức (QĐ 139/2024/YTCC-HD3).')

bl(d2, 'Tổng quan tài liệu của chúng tôi cho thấy vùng Đồng bằng sông Cửu Long (ĐBSCL) hầu như chưa có nghiên cứu về sức khỏe tâm thần thanh thiếu niên. So sánh: Cần Thơ lo âu 59,0% (nghiên cứu trước), Thanh Hóa 49,0% (Danh Lâm, 2022). Đây là nghiên cứu đầu tiên tại An Giang.', bold=True)

bl(d2, 'Dữ liệu bao gồm giới tính, khối lớp (10\u201312), học lực, kinh tế gia đình. Cỡ mẫu 366 HS, tỷ lệ nam/nữ cân bằng (50,5/49,5%).')

rb(d2, 'Kết quả nghiên cứu định lượng')
bl(d2, 'Tỷ lệ lo âu, trầm cảm, stress (N = 366, DASS-21):', bold=True)
tbl(d2,
    ['Tình trạng', 'Tỷ lệ tổng', 'Nhẹ', 'Vừa', 'Nặng', 'Rất nặng'],
    [['Lo âu', '61,2%', '9,3%', '24,0%', '12,6%', '15,3%'],
     ['Trầm cảm', '47,3%', '15,8%', '18,0%', '8,5%', '4,9%'],
     ['Stress', '38,0%', '12,8%', '16,4%', '11,5%', '4,9%']],
    widths=[2.5, 2.0, 1.5, 1.5, 1.5, 2.0])
d2.add_paragraph()

bl(d2, 'Đồng mắc (N = 366):', bold=True)
tbl(d2,
    ['Số triệu chứng', 'Tỷ lệ', 'Ghi chú'],
    [['Không có (0)', '28,1%', '\u2014'],
     ['Có 1 triệu chứng', '24,6%', '\u2014'],
     ['Có 2 triệu chứng', '20,0%', '\u2014'],
     ['Có cả 3 triệu chứng', '27,3%', 'Hơn 1/4 đồng mắc cả 3']],
    widths=[4.0, 2.5, 5.0])
d2.add_paragraph()

bh3(d2, 'So sánh liên bài : Tổng quan tài liệu cho thấy tỷ lệ lo âu bằng DASS-21 tại Việt Nam dao động rộng. Lo âu 61,2% tại An Giang nằm ở mức trung bình-cao: thấp hơn Bảo Quyên 2025 (86,2%) và Thảo Vi 2025 (65,8%), nhưng cao hơn Danh Lâm 2022 (49,0%) và Hoàng Trung Học 2025 (41,5% trong COVID). Nói cách khác, vùng ĐBSCL có tỷ lệ tương đương với các vùng khác khi dùng cùng công cụ DASS-21.')

rb(d2, 'Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3(d2, 'Phát hiện về đồng mắc và đặc thù ĐBSCL')
bl(d2, '*Đồng mắc cao.* 27,3% HS có cả 3 triệu chứng (lo âu + trầm cảm + stress) \u2014 gợi ý các rối loạn này chồng chéo mạnh, tương tự Bhardwaj 2020 (OR đồng mắc = 9,77\u201317,26). Điều này ủng hộ quan điểm sàng lọc cả 3 tình trạng đồng thời.')
bl(d2, '*Tỷ lệ giới cân bằng.* Nam 50,5% và nữ 49,5% \u2014 tốt hơn đáng kể so với Bảo Quyên 2025 (78% nữ), cho phép so sánh giới tính công bằng hơn. Tuy nhiên, bài không phân tích giới tính chi tiết.')
bl(d2, '*Lo âu chiếm tỷ lệ cao nhất.* Lo âu 61,2% > trầm cảm 47,3% > stress 38,0% \u2014 phù hợp với xu hướng quốc tế: lo âu là vấn đề phổ biến nhất ở VTN.')

rb(d2, 'Kết luận')
bl(d2, 'Dữ liệu của chúng tôi, cho thấy tỷ lệ lo âu 61,2%, trầm cảm 47,3% và stress 38,0% ở 366 HS THPT An Giang \u2014 nghiên cứu đầu tiên tại ĐBSCL, với 27,3% đồng mắc cả 3 triệu chứng, gợi ý rằng sức khỏe tâm thần VTN tại vùng ĐBSCL cần được quan tâm tương đương như các vùng đô thị lớn. Lo âu chiếm tỷ lệ cao nhất, ủng hộ ưu tiên sàng lọc và can thiệp lo âu tại trường học vùng nông thôn.', bold=True)

rh2(d2, 'Phản biện')
bl(d2, 'Điểm mạnh: NC đầu tiên ĐBSCL, tỷ lệ giới cân bằng, DASS-21 đã xác thực, có đạo đức, DOI. Tuy nhiên, chỉ 1 trường (Long Bình) \u2014 không đại diện. Chỉ thống kê mô tả, KHÔNG phân tích yếu tố liên quan. Bài ngắn (4 trang nội dung), thiếu thảo luận sâu. Không phân tích giới tính, khối lớp. Tạp chí Y học VN không lập chỉ mục quốc tế.')

rh2(d2, 'Hướng nghiên cứu tiếp theo')
bl(d2, 'Mở rộng nhiều trường An Giang và ĐBSCL. Phân tích hồi quy đa biến xác định yếu tố nguy cơ. So sánh ĐBSCL với đô thị (TPHCM, Hà Nội). Phân tích giới tính chi tiết. Can thiệp tại trường THPT vùng ĐBSCL.')

d2.add_paragraph()
p = d2.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50 Trung bình-Thấp. NC đầu tiên ĐBSCL, giới cân bằng, nhưng 1 trường, chỉ mô tả, bài ngắn, tạp chí không QT.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)
d2.save('VN18_AnGiang_2025.docx')
sys.stderr.write('VN18 OK\n')
