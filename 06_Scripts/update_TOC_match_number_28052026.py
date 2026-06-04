# -*- coding: utf-8 -*-
"""Match TOC row voi heading theo SO THU TU (1.1.1, 7.1, etc.) + cap nhat text + page.
Them dong moi cho headings KHONG match.
28/05/2026."""
import os, sys, io, re, shutil, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import win32com.client, pythoncom
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

# Backup
BACKUP = FILE.replace('.docx', '_BEFORE_TOC_REBUILD.docx')
shutil.copy(FILE, BACKUP)
print(f"Backup: {BACKUP}")


# ============================================================
# 1. Extract headings via Word COM
# ============================================================
print("\n[1] Extract headings + pages via Word COM...")
TEMP = os.path.join(ROOT, 'Luận án TS', '_temp_match.docx')
shutil.copy(FILE, TEMP)
pythoncom.CoInitialize()
word = None
headings = []
try:
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False; word.DisplayAlerts = False
    doc = word.Documents.Open(os.path.abspath(TEMP), ReadOnly=False)
    doc.Fields.Update(); doc.Repaginate()
    for level in [1, 2, 3, 4, 5]:
        for style_name in [f'Heading {level}', f'Tiêu đề {level}']:
            try:
                rng = doc.Content
                find = rng.Find
                find.ClearFormatting(); find.Style = doc.Styles(style_name)
                find.Text = ''; find.Forward = True; find.Wrap = 0; find.Format = True
                while find.Execute():
                    para = rng.Paragraphs.First
                    text = para.Range.Text.strip().replace('\r','').replace('\x07','')
                    if para.Range.Tables.Count > 0: rng.Collapse(0); continue
                    if not text or len(text) > 250: rng.Collapse(0); continue
                    page = para.Range.Information(3)
                    # Also get the paragraph order (rough position via Word's "Style" find sequence)
                    start_pos = para.Range.Start
                    headings.append({'level': level, 'text': text, 'page': page, 'pos': start_pos})
                    rng.Collapse(0)
            except: pass
    doc.Close(SaveChanges=False)
finally:
    if word:
        try: word.Quit()
        except: pass
    pythoncom.CoUninitialize()
    if os.path.exists(TEMP):
        try: os.remove(TEMP)
        except: pass

# Dedupe
seen = set(); unique = []
for h in headings:
    k = (h['text'], h['page'])
    if k not in seen:
        seen.add(k); unique.append(h)
headings = sorted(unique, key=lambda h: h.get('pos', h['page']*10000))
print(f"  Extracted {len(headings)} headings")


# ============================================================
# 2. Match TOC rows to headings (PRIMARY by numbering prefix)
# ============================================================
def extract_num(text):
    """Get numbering prefix like '1.1.1' or '7.1' or None."""
    m = re.match(r'^(\d+(?:\.\d+)*)', text.strip())
    return m.group(1) if m else None

def normalize_text(s):
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'^[\d\.]+\s*', '', s)
    return s.strip(' .')


# Build heading lookups
by_num = {}
by_text = {}
for h in headings:
    num = extract_num(h['text'])
    if num:
        by_num.setdefault(num, []).append(h)
    norm = normalize_text(h['text'])
    if norm and len(norm) > 4:
        by_text[norm] = h


# Open doc, find TOC table
d = Document(FILE)
toc_table = None
toc_table_idx = None
for ti, tb in enumerate(d.tables):
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb
        toc_table_idx = ti
        break

print(f"\n[2] TOC table = Table {toc_table_idx} ({len(toc_table.rows)} rows)")


# ============================================================
# 3. Update existing TOC rows by NUMBER match
# ============================================================
print("\n[3] Update TOC rows by number-prefix match...")
matched_headings = set()  # which heading was matched
updates = []
for ri, row in enumerate(toc_table.rows):
    name = row.cells[0].text.strip()
    old_page = row.cells[1].text.strip() if len(row.cells) > 1 else ''
    num = extract_num(name)
    matched_h = None

    if num and num in by_num:
        # Take first matching heading
        candidates = by_num[num]
        for h in candidates:
            if (h['text'], h['page']) in matched_headings: continue
            matched_h = h
            matched_headings.add((h['text'], h['page']))
            break

    if matched_h is None:
        # Try text match
        norm = normalize_text(name)
        if norm in by_text:
            h = by_text[norm]
            if (h['text'], h['page']) not in matched_headings:
                matched_h = h
                matched_headings.add((h['text'], h['page']))

    if matched_h:
        new_page = str(matched_h['page'])
        if new_page != old_page:
            cell = row.cells[1]
            font_name = 'Times New Roman'; font_size = Pt(13)
            if cell.paragraphs and cell.paragraphs[0].runs:
                old_r = cell.paragraphs[0].runs[0]
                if old_r.font.name: font_name = old_r.font.name
                if old_r.font.size: font_size = old_r.font.size
            for p in cell.paragraphs:
                for r in p.runs:
                    r.text = ''
            cell.paragraphs[0].text = ''
            new_run = cell.paragraphs[0].add_run(new_page)
            new_run.font.name = font_name
            new_run.font.size = font_size
            updates.append((ri, name, old_page, new_page))

print(f"  Updated {len(updates)} rows by number match")

# Headings NOT matched = need to add
not_matched_headings = []
for h in headings:
    key = (h['text'], h['page'])
    if key in matched_headings: continue
    # Skip cover-type
    upper = h['text'].upper().strip()
    if upper in ('MỤC LỤC', 'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
                 'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
                 'LỜI CAM ĐOAN', 'LỜI CẢM ƠN'):
        continue
    not_matched_headings.append(h)

print(f"\n[4] {len(not_matched_headings)} headings KHONG match - se THEM vao TOC")
for h in not_matched_headings[:30]:
    print(f"  L{h['level']} p{h['page']}: {h['text'][:80]}")
if len(not_matched_headings) > 30:
    print(f"  ... va {len(not_matched_headings)-30} muc khac")

d.save(FILE)
print(f"\nSaved updates: {FILE}")


# Save not-matched for next step
NM_FILE = os.path.join(ROOT, '06_Scripts', '_not_matched_headings.json')
with open(NM_FILE, 'w', encoding='utf-8') as f:
    # Strip non-serializable 'pos'
    out = [{'level': h['level'], 'text': h['text'], 'page': h['page']} for h in not_matched_headings]
    json.dump(out, f, ensure_ascii=False, indent=2)
print(f"Saved unmatched headings: {NM_FILE}")
