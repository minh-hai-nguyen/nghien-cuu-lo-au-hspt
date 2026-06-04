"""DOC 1: Tinh ro 'khong khop giua 85,4% va r²=0,176' trong VN004 Nguyen Thi Van.
Pure mathematical reasoning + interpretation. KHONG bia so lieu.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/VN004_85_4_khong_khop_giai_thich.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE
    return p

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('VÌ SAO "85,4% GIẢI THÍCH" KHÔNG KHỚP\nVỚI r² = 0,1764 (TỪ r = 0,42)?\n— Phân tích toán học chi tiết VN004 Nguyễn Thị Vân (2020) —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Tính cho thầy rõ "không khớp" là như thế nào nhé. Nguyễn Thị Vân '
    'tỷ lệ "85,4%" cần được làm rõ về cách tính (không khớp với r² = '
    '0,176 từ r = 0,42).', italic=True
)

# 1. Hai con số trong bài VN004
H('1. Hai con số trong VN004', level=2, color=NAVY)
para(
    'Trong báo cáo VN004 Nguyễn Thị Vân (2020) về "Mức độ lo âu '
    'của học sinh trung học phổ thông thành phố Hồ Chí Minh", nhóm '
    'yếu tố "bản thân học sinh" được công bố hai số liệu trên cùng '
    'một mục.', indent=False
)
add_table(
    ['Chỉ số trong VN004', 'Giá trị tác giả công bố'],
    [
        ['Hệ số tương quan r (nhóm yếu tố bản thân HS với lo âu)', 'r = 0,42'],
        ['"Tỷ lệ giải thích" của nhóm yếu tố bản thân HS', '85,4%'],
    ]
)
para('')
para(
    'Câu hỏi đặt ra: hai con số 0,42 và 85,4% có khớp nhau theo lý '
    'thuyết thống kê không? Để trả lời, cần đi qua bốn bước toán học.'
)

# 2. Phép tính r² từ r = 0,42
H('2. Phép tính r² từ hệ số tương quan r = 0,42', level=2, color=NAVY)
para(
    'Trong thống kê tuyến tính, nếu r là hệ số tương quan Pearson '
    'giữa hai biến X và Y, thì r² (hệ số XÁC ĐỊNH — coefficient of '
    'determination) là TỶ LỆ PHƯƠNG SAI của Y được giải thích bởi '
    'X qua hồi quy tuyến tính bivariate.', indent=False
)
para('Áp dụng cho VN004:', bold=True)
para(
    'r = 0,42 → r² = 0,42 × 0,42 = 0,1764', size=14, bold=True, indent=False
)
para(
    'Tức là 17,64% phương sai của lo âu được giải thích bởi mối quan '
    'hệ tuyến tính với nhóm yếu tố bản thân học sinh.'
)

# 3. So sánh hai con số
H('3. So sánh hai con số — gap rất lớn', level=2, color=NAVY)
caption('Bảng 1. Hai con số được công bố trong VN004 và phép tính so sánh')
add_table(
    ['Đại lượng', 'Giá trị', 'Ý nghĩa thống kê'],
    [
        ['r² tính từ r = 0,42',
         '0,1764 = 17,64%',
         '% phương sai Y giải thích bởi X qua hồi quy tuyến tính 1 biến'],
        ['"Tỷ lệ giải thích" tác giả công bố',
         '0,854 = 85,4%',
         'Không nêu rõ phương pháp tính'],
        ['Chênh lệch tuyệt đối',
         '85,4% − 17,64% = 67,76 điểm phần trăm',
         'Gap rất lớn — không thể bỏ qua'],
        ['Tỷ số',
         '85,4% / 17,64% ≈ 4,84 lần',
         '85,4% lớn HƠN gần 5 lần so với r²'],
        ['Hệ số r tương đương để có 85,4%',
         '√0,854 ≈ 0,924',
         'Cần r = 0,924 để có r² = 85,4% — gần gấp đôi 0,42'],
    ]
)
para('')
para(
    'Nói cách khác, để có 85,4% giải thích từ một mối tương quan '
    'tuyến tính bivariate, hệ số r phải bằng √0,854 ≈ 0,924 — gần '
    'GẤP ĐÔI giá trị 0,42 mà tác giả công bố. Hai con số này KHÔNG '
    'thể đồng thời đúng nếu cùng đo "tỷ lệ giải thích" của một mối '
    'quan hệ tuyến tính bivariate.'
)

# 4. Bốn khả năng cách hiểu 85,4%
H('4. Bốn khả năng giải thích con số 85,4%', level=2, color=NAVY)
para(
    'Vì r² = 17,64% và 85,4% không thể cùng là một đại lượng, con số '
    '85,4% trong báo cáo VN004 phải được hiểu theo MỘT trong bốn cách '
    'sau — mỗi cách có hàm ý khác nhau về giá trị khoa học của bài.', indent=False
)

H('Khả năng 1 — 85,4% là R² của hồi quy ĐA BIẾN', level=3)
para(
    'Trong hồi quy bội (multiple regression) với nhiều biến predictor '
    'cùng nhóm "bản thân học sinh" (ví dụ: tự ti, lo sợ, thất vọng, '
    'ít giao tiếp...), hệ số xác định R² (chữ R hoa) tổng hợp tác '
    'động của TẤT CẢ predictor — KHÔNG phải r² (chữ r thường) của '
    'một biến đơn.', indent=False
)
para(
    'Phép toán: nếu cụm có 5–7 biến predictor cùng tương quan vừa '
    'với lo âu thì R² hồi quy bội có thể đạt 80–90%. Tuy nhiên, '
    'trong trường hợp này, tác giả phải báo cáo: (a) ma trận tương '
    'quan giữa các predictor; (b) hệ số β chuẩn hóa của từng biến; '
    '(c) F-test cho R²; (d) kiểm tra đa cộng tuyến (VIF). KHÔNG có '
    'các thông tin này trong bản báo cáo VN004.'
)
para(
    'Hệ quả: nếu khả năng 1 đúng, BÁO CÁO VN004 THIẾU minh bạch '
    'phương pháp luận — vi phạm STROBE (von Elm và cộng sự, 2007).'
)

H('Khả năng 2 — 85,4% là TỶ LỆ HỌC SINH có biểu hiện', level=3)
para(
    'Đây là giải thích đơn giản nhất. 85,4% có thể là tỷ lệ học '
    'sinh trong mẫu CÓ ÍT NHẤT MỘT biểu hiện thuộc nhóm "bản thân '
    'HS" (ít giao tiếp, thất vọng, lo sợ...). Đây là TỶ LỆ HIỆN '
    'MẮC (prevalence rate) — KHÔNG có quan hệ nào với r².', indent=False
)
para(
    'Hệ quả: nếu khả năng 2 đúng, tác giả đã NHẦM thuật ngữ. '
    '"Tỷ lệ % học sinh có biểu hiện" KHÁC HOÀN TOÀN với "tỷ lệ % '
    'phương sai giải thích". Hai khái niệm có ý nghĩa thống kê '
    'TRÁI ngược nhau:'
)
para(
    '• Tỷ lệ % HS có biểu hiện = mô tả mẫu, KHÔNG nói về quan hệ.\n'
    '• Tỷ lệ % phương sai giải thích = đo cường độ quan hệ giữa X và Y.', indent=False
)

H('Khả năng 3 — 85,4% là tổng phản hồi định tính từ phỏng vấn sâu', level=3)
para(
    'VN004 sử dụng thiết kế hai pha — sàng lọc 558 HS bằng STAI rồi '
    'phỏng vấn sâu 90 HS. 85,4% có thể là tỷ lệ học sinh trong 90 HS '
    'phỏng vấn nhắc đến yếu tố thuộc nhóm "bản thân" trong câu trả '
    'lời của mình. Tức là 85,4% × 90 ≈ 77 học sinh có nhắc đến.', indent=False
)
para(
    'Hệ quả: nếu khả năng 3 đúng, đây là DỮ LIỆU ĐỊNH TÍNH — không '
    'có ý nghĩa thống kê định lượng và không thể trình bày như tỷ '
    'lệ "giải thích" trong bài báo định lượng.'
)

H('Khả năng 4 — 85,4% là sai sót đánh máy hoặc dịch chuyển dấu phẩy', level=3)
para(
    'Khả năng kỹ thuật: 85,4% có thể là 8,54% (dịch sai dấu phẩy) — '
    'tương đương r² = 0,0854 → r = 0,292. Hoặc 85,4% là tỷ số khác '
    'bị nhầm thành "tỷ lệ giải thích". Đây là khả năng ÍT có nhất '
    'nhưng cần loại trừ.', indent=False
)

# 5. Bảng tổng hợp 4 khả năng
H('5. Bảng tổng hợp bốn khả năng', level=2, color=NAVY)
caption('Bảng 2. Bốn khả năng giải thích con số 85,4% trong VN004')
add_table(
    ['Khả năng', 'Bản chất 85,4%', 'Phép toán kiểm chứng', 'Hệ quả phương pháp luận'],
    [
        ['1. R² hồi quy đa biến',
         '% phương sai Y giải thích bởi NHIỀU predictor cùng nhóm',
         'R² ≠ r² của 1 biến — toán học HỢP LỆ',
         'Tác giả thiếu báo cáo β chuẩn hóa, F-test, VIF — vi phạm STROBE'],
        ['2. Tỷ lệ HS có biểu hiện',
         'Prevalence rate trong mẫu',
         '85,4% × 558 ≈ 477 HS có biểu hiện',
         'NHẦM thuật ngữ — tỷ lệ HS ≠ tỷ lệ phương sai giải thích'],
        ['3. Phản hồi định tính phỏng vấn',
         '85,4% × 90 ≈ 77 HS phỏng vấn nhắc đến',
         'Dữ liệu định tính, không có ý nghĩa định lượng',
         'KHÔNG nên trình bày như con số "giải thích"'],
        ['4. Sai sót đánh máy',
         'Có thể là 8,54% (r ≈ 0,292)',
         '8,54% gần với r² của r ≈ 0,3',
         'Cần erratum — sửa lại bản gốc'],
    ]
)

# 6. Kết luận
H('6. Kết luận', level=2, color=NAVY)
para(
    'Hai con số 0,42 và 85,4% trong VN004 KHÔNG thể đồng thời đúng '
    'nếu cùng đo "tỷ lệ giải thích" của một tương quan tuyến tính '
    'bivariate. Phép toán cho thấy:', indent=False
)
para(
    '• Nếu r = 0,42 → r² = 0,1764 = 17,64% → 85,4% phải có ý nghĩa '
    'KHÁC.\n'
    '• Nếu 85,4% là r² đúng → cần r = 0,924 (gần gấp đôi 0,42 đã '
    'công bố) → mâu thuẫn nội tại.\n'
    '• Chênh lệch 67,76 điểm phần trăm là QUÁ LỚN để có thể bỏ qua.', indent=False
)
para(
    'Khả năng cao nhất là 85,4% là R² của hồi quy đa biến (Khả năng '
    '1) — toán học hợp lệ nhưng tác giả thiếu báo cáo các thống kê '
    'kèm theo (β chuẩn hóa cho từng biến, F-test cho mô hình, VIF '
    'cho đa cộng tuyến). Trước khi trích dẫn 85,4% vào báo cáo CTH, '
    'thầy nên LIÊN HỆ tác giả gốc để xác nhận phương pháp tính, hoặc '
    'KHÔNG dùng con số này — chỉ dùng r = 0,42 (đã công bố rõ).'
)

# 7. CÂU TRẢ LỜI tô xanh
H('7. CÂU TRẢ LỜI', level=2, color=NAVY)
blue_run('Tóm gọn không khớp giữa 85,4% và r² = 0,1764:', bold=True)
blue_run(
    '• r = 0,42 (Pearson) → r² = 0,42² = 0,1764 = 17,64%. Đây là '
    'TỶ LỆ PHƯƠNG SAI Y giải thích bởi X qua hồi quy tuyến tính '
    'bivariate.'
)
blue_run(
    '• Tác giả công bố "85,4% giải thích" — gấp 4,84 lần r² lý '
    'thuyết, chênh lệch 67,76 điểm phần trăm.'
)
blue_run(
    '• Để có 85,4% từ r² thực, cần r = √0,854 ≈ 0,924 — gần gấp '
    'đôi 0,42 đã công bố. Mâu thuẫn nội tại.'
)
blue_run(
    '• Bốn khả năng giải thích: (1) R² hồi quy đa biến — toán '
    'hợp lệ nhưng thiếu β/F/VIF; (2) tỷ lệ % HS có biểu hiện — '
    'nhầm thuật ngữ; (3) tỷ lệ định tính từ phỏng vấn — không '
    'thống kê; (4) sai sót đánh máy.'
)
blue_run(
    '• Khuyến nghị: KHÔNG trích con số 85,4% vào báo cáo CTH '
    'của thầy — chỉ trích r = 0,42 (đã công bố rõ) và ghi chú '
    'rằng "tỷ lệ giải thích 85,4% trong báo cáo gốc cần làm rõ '
    'phương pháp tính".'
)

# 8. Phụ lục
H('8. Phụ lục — Công thức và tham chiếu', level=2, color=NAVY)
para('Công thức 1 — Hệ số xác định:', bold=True, indent=False)
para('r² = 1 − (SS_res / SS_tot) = (Cov(X,Y))² / (Var(X) × Var(Y))', italic=True, indent=False)
para('Trong tương quan tuyến tính bivariate, r² = (Pearson r)².', indent=False)

para('')
para('Công thức 2 — Hệ số xác định bội:', bold=True, indent=False)
para('R² = SS_reg / SS_tot (với nhiều predictor)', italic=True, indent=False)
para('R² ≥ r² của bất kỳ biến đơn nào trong mô hình.', indent=False)

para('')
para('Tham chiếu lý thuyết:', bold=True, indent=False)
refs = [
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates. [Quy ước cường độ r: 0,10 = nhỏ; 0,30 = trung bình; 0,50 = lớn.]',
    'Field, A. (2018). Discovering statistics using IBM SPSS Statistics (5th ed.). Sage.',
    'Nguyễn, T. V. (2020). Mức độ lo âu của học sinh trung học phổ thông thành phố Hồ Chí Minh. [VN004 trong cơ sở dữ liệu dự án.]',
    'von Elm, E., Altman, D. G., Egger, M., Pocock, S. J., Gøtzsche, P. C., & Vandenbroucke, J. P. (2007). The Strengthening the Reporting of Observational Studies in Epidemiology (STROBE) statement. Annals of Internal Medicine, 147(8), 573–577. https://doi.org/10.7326/0003-4819-147-8-200710160-00010',
]
for r in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    run = p.add_run(r); run.font.size = Pt(11)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
