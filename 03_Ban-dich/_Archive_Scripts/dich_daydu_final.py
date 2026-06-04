# -*- coding: utf-8 -*-
"""Dịch đầy đủ 9 bài QT26-34 + bổ sung 5 bài VN16-20 — FINAL"""
import sys, os, fitz
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

PDF = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '02_Papers-goc', 'The-gioi-moi')
VN_PDF = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', '02_Papers-goc', 'Viet-Nam')
OUT = os.path.dirname(os.path.abspath(__file__))
CHARTS = os.path.join(OUT, 'Charts')

def save(doc, name):
    p = os.path.join(OUT, name)
    doc.save(p)
    import docx as dx
    d = dx.Document(p)
    t = '\n'.join([x.text for x in d.paragraphs])
    print(f'  {name}: {len(t)} chars, {len(d.tables)} tables')

def read_pdf_pages(pdf_path, pages):
    """Read text from specific pages"""
    d = fitz.open(pdf_path)
    texts = []
    for p in pages:
        if p < len(d):
            texts.append(d[p].get_text())
    return '\n'.join(texts)

def crop_fig(pdf_path, page_idx, y1_frac, y2_frac, out_name):
    d = fitz.open(pdf_path)
    page = d[page_idx]
    pix = page.get_pixmap(dpi=200, clip=fitz.Rect(0, page.rect.height*y1_frac, page.rect.width, page.rect.height*y2_frac))
    path = os.path.join(CHARTS, out_name)
    pix.save(path)
    return path

# =====================================================================
# QT26 — UK NHS Parliament 2024 (46p) — COMPLETE
# =====================================================================
def dich_QT26():
    print('QT26 UK NHS (full)...')
    doc = create_doc()
    add_p(doc, 'Link: https://commonslibrary.parliament.uk/research-briefings/sn06988/', size=10)
    add_heading(doc, 'Thống kê sức khỏe tâm thần: Tỷ lệ, dịch vụ và tài trợ tại Anh', 1)
    h = doc.add_paragraph(); r = h.add_run('Mental Health Statistics: Prevalence, Services and Funding in England'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Mental Health Statistics: Prevalence, Services and Funding in England'),
        ('Tác giả', 'Carl Baker, Esme Kirk-Wade'),
        ('Cơ quan', 'UK Parliament — House of Commons Library'),
        ('Xuất bản', '1 March 2024, Briefing Paper SN06988, 46 trang'),
        ('Loại tài liệu', 'Báo cáo thống kê chính sách (briefing paper)'),
        ('Phạm vi', 'England (Anh quốc), dữ liệu NHS'),
    ])
    add_page_ref(doc, '1–46', 'UK Parliament Briefing', 'SN06988, March 2024')

    # TÓM TẮT
    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Báo cáo tổng hợp toàn diện thống kê SKTT tại Anh: tỷ lệ rối loạn, dịch vụ NHS, chi tiêu, can thiệp. Dữ liệu chính: Ước tính 1/6 người trưởng thành Anh đã trải qua rối loạn tâm thần phổ biến (trầm cảm, lo âu) trong tuần qua. Ở trẻ em: tỷ lệ rối loạn tâm thần có thể xác định ở trẻ 8–16 tuổi tăng từ 12% (2017) lên 20% (2023). Ở nhóm 17–19: tăng từ 10% lên 23%. Chi tiêu SKTT NHS: £16 tỷ/năm (14% ngân sách NHS địa phương).')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    # ĐÁNH GIÁ NHANH
    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['UK Parliament — nguồn dữ liệu chính sách uy tín nhất.', 'Tỷ lệ RLTT trẻ 8–16: 12% (2017) → 20% (2023) — tăng gấp rưỡi.', 'Nhóm 17–19: 10% → 23% — tăng hơn gấp đôi, mạnh nhất.', 'Nữ 17–19: 32% có RLTT có thể — cao gần GẤP ĐÔI nam (15%).', 'Dịch vụ: 3,58 triệu người tiếp xúc NHS SKTT/năm; 1,76 triệu giới thiệu TTAD.', 'Chi tiêu: £16 tỷ/năm — 14% ngân sách NHS. VN: chưa có số liệu tương đương.', 'TTAD (Talking Therapies): 66,5% cải thiện, 49,9% phục hồi — tiêu chuẩn 50%.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Chỉ England — không đại diện Scotland, Wales, Northern Ireland.', 'Briefing paper — tổng hợp, không có peer review.', 'NHS phổ quát miễn phí — bối cảnh khác VN hoàn toàn.', 'Dữ liệu 2023 là ước tính — "probable" disorder, không phải chẩn đoán chắc chắn.']:
        add_p(doc, f'• {b}')

    # 1. TỶ LỆ NGƯỜI LỚN
    add_page_ref(doc, '5–13', 'UK Parliament Briefing', 'SN06988')
    add_heading(doc, '1. TỶ LỆ RỐI LOẠN TÂM THẦN Ở NGƯỜI LỚN', 2)
    add_p(doc, 'Khảo sát SKTT Người lớn Anh (Adult Psychiatric Morbidity Survey) thực hiện năm 1993, 2000, 2007, 2014. Phát hiện chính: 1/6 người trưởng thành (17%) trải qua rối loạn tâm thần phổ biến (CMD — Common Mental Disorder) như trầm cảm hoặc lo âu trong tuần qua. Tỷ lệ tương đối ổn định qua các năm khảo sát.')
    add_p(doc, 'Rối loạn lo âu tổng quát (GAD) là CMD phổ biến nhất: 6,0% (2014). Tiếp theo: trầm cảm 3,3%, phobia 2,4%, OCD 1,3%, rối loạn hoảng sợ 0,6%.')
    add_p(doc, 'Khác biệt giới: Nữ có tỷ lệ CMD cao hơn nam (1/5 nữ vs 1/8 nam). Khoảng cách giới mở rộng từ 1993 đến 2014.')

    # 2. TRẺ EM VÀ THANH THIẾU NIÊN
    add_page_ref(doc, '14–16', 'UK Parliament Briefing', 'SN06988')
    add_heading(doc, '2. SỨC KHỎE TÂM THẦN TRẺ EM VÀ THANH THIẾU NIÊN', 2)
    add_p(doc, 'Khảo sát SKTT Trẻ em và Thanh niên NHS (NHS Digital surveys) thực hiện 2017, 2020, 2021, 2022, 2023. Đây là nguồn dữ liệu chính cho SKTT VTN tại Anh.')

    add_heading(doc, 'Bảng 1. Tỷ lệ rối loạn tâm thần "có thể" ở trẻ em/VTN Anh', 3)
    add_table(doc,
        ['Nhóm tuổi', '2017', '2020', '2021', '2022', '2023', 'Xu hướng'],
        [['8–10 tuổi', '—', '—', '—', '—', '17%', '(Không so sánh trực tiếp)'],
         ['8–16 tuổi', '12%', '—', '—', '—', '20%', '↑ Tăng gấp rưỡi'],
         ['11–16 tuổi', '—', '—', '—', '—', '21%', '↑ Cao hơn 2017 đáng kể'],
         ['17–19 tuổi', '10%', '18%', '17%', '26%', '23%', '↑ TĂNG HƠN GẤP ĐÔI'],
         ['Nữ 17–19', '—', '—', '—', '—', '32%', 'CAO NHẤT — gấp đôi nam'],
         ['Nam 17–19', '—', '—', '—', '—', '15%', '']],
        widths=[2.5, 1.5, 1.5, 1.5, 1.5, 1.5, 3.5])
    add_p(doc, 'Ghi chú: "Probable mental disorder" dựa trên SDQ. Năm 2020 = COVID bùng phát. Tỷ lệ tăng giữa 2017–2020, ổn định 2020–2021, tăng lại 2021–2022, giảm nhẹ 2022–2023 nhưng vẫn cao hơn trước đại dịch đáng kể.', size=9, italic=True)

    add_p(doc, 'Giới tính và tuổi:', bold=True)
    add_p(doc, '• Trẻ 8–10: NAM nhiều hơn nữ (20% vs 14%) — mô hình ngoại hóa.')
    add_p(doc, '• VTN 17–19: NỮ nhiều hơn nam (32% vs 15%) — mô hình nội hóa. Khoảng cách GẤP ĐÔI.')
    add_p(doc, '• Dân tộc: Trẻ nhóm Trắng có tỷ lệ cao hơn (20%) so với nhóm thiểu số (10%) năm 2021 — NGƯỢC với người lớn.')
    add_p(doc, '• SEND (Nhu cầu Giáo dục Đặc biệt): 57% có RLTT có thể — gấp 4 lần trẻ không SEND (13%).')

    # 3. DỊCH VỤ NHS
    add_page_ref(doc, '19–27', 'UK Parliament Briefing', 'SN06988')
    add_heading(doc, '3. DỊCH VỤ SKTT NHS', 2)

    add_heading(doc, 'Bảng 2. Dịch vụ SKTT NHS England', 3)
    add_table(doc,
        ['Chỉ số', 'Giá trị', 'Năm', 'Ghi chú'],
        [['Người tiếp xúc SKTT NHS', '3,58 triệu', '2022/23', '6% dân số; tăng 24% so với 2019/20'],
         ['Nhập viện nội trú', '92.000', '2022/23', '<3% tổng tiếp xúc'],
         ['VTN 11–15 tiếp xúc NHS', '16,7% dân số nhóm tuổi', '2022/23', 'Cao nhất trong các nhóm tuổi'],
         ['VTN 16–19 tiếp xúc NHS', '14,6%', '2022/23', ''],
         ['Giới thiệu TTAD', '1,76 triệu', '2022/23', 'Giảm 2,9% so với 2021/22'],
         ['Hoàn thành điều trị TTAD', '672.000', '2022/23', '38% số giới thiệu'],
         ['TTAD cải thiện', '66,5%', '2022/23', 'Giảm từ 68,3% (2020/21)'],
         ['TTAD phục hồi', '49,9%', '2022/23', 'Mục tiêu 50% — gần đạt'],
         ['Thời gian chờ trung vị', '45 ngày', 'Q4 2023', 'Dao động 11–121 ngày theo vùng'],
         ['Thời gian chờ P90', '251 ngày', 'Q4 2023', 'Dài nhất: 1.087 ngày (Cheshire)']],
        widths=[4.0, 2.5, 2.0, 5.0])

    add_p(doc, 'TTAD (Talking Therapies for Anxiety and Depression):', bold=True)
    add_p(doc, 'Chương trình liệu pháp tâm lý NHS, ra mắt 2008 (trước là IAPT). Bao gồm CBT, tư vấn, hỗ trợ tự lực. Người lớn tuổi lao động có CMD. Tự giới thiệu hoặc GP giới thiệu. 1,76 triệu giới thiệu/năm. Cải thiện 66,5%, phục hồi 49,9%. GAD có tỷ lệ cải thiện cao nhất (70,7%).')

    add_p(doc, 'Người trẻ kém hơn:', bold=True)
    add_p(doc, '• VTN dưới 18: chỉ 57% vào điều trị, 20% hoàn thành (thấp hơn nhiều so với người lớn 70%/40%).')
    add_p(doc, '• Phục hồi thấp hơn ở nhóm 25 trở xuống — không đạt mục tiêu 50%.')
    add_p(doc, '• Nữ nhiều khả năng cải thiện hơn, nhưng nam nhiều khả năng phục hồi hơn.')

    # 4. CHI TIÊU
    add_page_ref(doc, '30–40', 'UK Parliament Briefing', 'SN06988')
    add_heading(doc, '4. CHI TIÊU SKTT NHS', 2)

    add_heading(doc, 'Bảng 3. Chi tiêu SKTT NHS England', 3)
    add_table(doc,
        ['Chỉ số', 'Giá trị', 'Ghi chú'],
        [['Tổng chi tiêu SKTT NHS', '£16 tỷ/năm', '14% ngân sách NHS địa phương (2022/23)'],
         ['Tăng so với 2016/17', '+£4,7 tỷ', 'Tăng từ £11,6 tỷ'],
         ['Chi tiêu SKTT CYP', '£1 tỷ+', 'Ước tính — khó tách riêng'],
         ['Chi tiêu bình quân/người', '~£280/năm', 'Dao động theo vùng'],
         ['So sánh với VN', 'CHƯA CÓ số liệu VN', 'VN: 8,4% tiếp cận vs NHS phổ quát']],
        widths=[4.0, 3.0, 6.0])

    # 5. COVID VÀ GIÁ SINH HOẠT
    add_heading(doc, '5. TRẦM CẢM, COVID-19, VÀ GIÁ SINH HOẠT', 2)
    add_p(doc, 'ONS theo dõi trầm cảm người lớn GB trong và sau COVID. Trầm cảm trung bình–nặng: 10% trước COVID → 21% đỉnh (1–3/2021) → 16% (9–10/2022). Vẫn cao hơn trước đại dịch.')
    add_p(doc, 'Chi phí sinh hoạt liên quan: 24% người khó trả hóa đơn năng lượng có trầm cảm (vs 9% người dễ trả). 27% người thuê nhà có trầm cảm (vs 10% sở hữu nhà).')

    # 6. KẾT LUẬN + PHẢN BIỆN
    add_heading(doc, '6. KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu NHS England cho thấy RLTT trẻ em/VTN tăng đáng kể: 12% → 20% (8–16 tuổi), 10% → 23% (17–19 tuổi), nữ 17–19 lên 32%. NHS chi £16 tỷ/năm nhưng vẫn không đáp ứng đủ nhu cầu (thời gian chờ dài, phục hồi chỉ 49,9%). So với VN (8,4% tiếp cận, chưa có số liệu chi tiêu) — khoảng cách khổng lồ.')

    # TLTK
    add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
    add_p(doc, 'Baker, C. & Kirk-Wade, E. (2024). Mental Health Statistics: Prevalence, Services and Funding in England. House of Commons Library, SN06988.', size=10)
    add_p(doc, 'NHS Digital (2023). Mental Health of Children and Young People in England 2023.', size=10)
    add_p(doc, 'ONS (2022). Cost of living and depression in adults, Great Britain.', size=10)
    add_p(doc, '(Xem đầy đủ trong bài gốc — 30+ nguồn)', size=10, italic=True)

    # VIẾT TẮT
    add_abbreviation_table(doc, [
        ('NHS', 'National Health Service — Dịch vụ Y tế Quốc gia Anh'),
        ('CMD', 'Common Mental Disorder — Rối loạn Tâm thần Phổ biến'),
        ('GAD', 'Generalised Anxiety Disorder — Rối loạn Lo âu Tổng quát'),
        ('SDQ', 'Strengths and Difficulties Questionnaire — Bảng hỏi Điểm mạnh Khó khăn'),
        ('TTAD', 'Talking Therapies for Anxiety and Depression (trước: IAPT)'),
        ('IAPT', 'Improving Access to Psychological Therapies'),
        ('CBT', 'Cognitive-Behavioural Therapy — Liệu pháp Nhận thức-Hành vi'),
        ('SEND', 'Special Educational Needs and Disabilities — Nhu cầu Giáo dục Đặc biệt'),
        ('ONS', 'Office for National Statistics — Cục Thống kê Quốc gia Anh'),
        ('PHQ8', 'Patient Health Questionnaire 8-item'),
        ('CYP', 'Children and Young People — Trẻ em và Thanh niên'),
        ('CYPMHS', "Children and Young People's Mental Health Services"),
    ])

    # PHẢN BIỆN
    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in [
        'UK Parliament — nguồn chính thức, dữ liệu NHS đáng tin cậy nhất.',
        'Xu hướng dài hạn 2017–2023 với 5 thời điểm — cho thấy tác động COVID rõ ràng.',
        'Bao gồm tỷ lệ + dịch vụ + chi tiêu + can thiệp — bức tranh toàn diện nhất trong Đề tài.',
        'Phân tầng theo tuổi, giới, dân tộc, SEND, vùng — rất chi tiết.',
        'TTAD (CBT quốc gia): 1,76 triệu giới thiệu/năm — chương trình can thiệp tâm lý lớn nhất thế giới. Phù hợp để học hỏi mô hình cho VN (Zugman et al. 2024, AJP: CBT hàng đầu).',
    ]:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế chi tiết:', bold=True)
    for s in [
        'Briefing paper — tổng hợp, KHÔNG có peer review. Dữ liệu từ nhiều nguồn khác nhau.',
        'Chỉ England — Scotland, Wales, NI có hệ thống riêng, dữ liệu khác.',
        'NHS phổ quát miễn phí — bối cảnh hoàn toàn khác VN (y tế tư nhân, bảo hiểm hạn chế, thiếu nhân lực SKTT).',
        '"Probable mental disorder" (SDQ) — sàng lọc, không chẩn đoán. So với V-NAMHS 2022 (DISC-5 chẩn đoán: 2,3%) — khác phương pháp nên không so sánh trực tiếp tỷ lệ.',
        'Không có dữ liệu lo âu riêng cho VTN — chỉ RLTT chung. JAACAP 2024 (QT23) tách lo âu riêng.',
        'Thời gian chờ dao động cực lớn (11–1.087 ngày) — phản ánh bất bình đẳng vùng ngay cả trong NHS phổ quát.',
    ]:
        add_red(doc, f'• {s}')
    add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
    for s in [
        'VN CHƯA CÓ báo cáo thống kê tương đương — cần tổng hợp tỷ lệ + dịch vụ + chi tiêu SKTT VTN. V-NAMHS 2022 là bước đầu nhưng chưa đủ chi tiết.',
        'So sánh chi tiêu SKTT: UK £16 tỷ vs VN chưa có số liệu. Ước tính chi tiêu/đầu người để đánh giá khoảng cách.',
        'Mô hình TTAD/IAPT (CBT quốc gia) — có thể áp dụng cho VN? Cần đánh giá tính khả thi (nhân lực, chi phí, văn hóa). BMC NMA 2025 (QT29) xác nhận CBT hiệu quả hàng đầu.',
        'SEND có tỷ lệ RLTT 57% — VN có dữ liệu tương tự cho trẻ khuyết tật không? Ngô Anh Vinh 2024 (VN15): DTTS 54,4% — mô hình tương tự.',
        'VTN dưới 18 hoàn thành điều trị thấp (20%) — rào cản gì? So với VN: chỉ 8,4% tiếp cận → rào cản còn lớn hơn nhiều.',
    ]:
        add_red(doc, f'• {s}')
    save(doc, '26_UK_NHS_2025_Parliament.docx')

# RUN
dich_QT26()
print('QT26 DONE!')
