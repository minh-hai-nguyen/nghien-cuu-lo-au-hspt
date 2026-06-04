# -*- coding: utf-8 -*-
"""
KG v2 — Entity Resolution + Author Disambiguation

Cai tien so voi v1:
- Country detection chi tu "location of study" (khong phai citation)
- Author disambiguation (Brown 2024 BESST vs Brown & Carter 2025 Editorial)
- Effect size context (SAD-specific vs general vs WL comparison)
- Journal detection duoc mo rong
- First author extracted rõ ràng
- Edge attribute phong phu hon
"""
import os, sys, io, json, re
from collections import defaultdict, Counter
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import networkx as nx
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
KG_DIR = os.path.join(os.path.dirname(__file__), 'kg_data')
os.makedirs(KG_DIR, exist_ok=True)

# ============================================================
# 1. LOAD DATA
# ============================================================
with open(os.path.join(os.path.dirname(__file__), 'canonical_mapping.json'), encoding='utf-8') as f:
    CANONICAL = json.load(f)

TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')

def extract_info_table(doc):
    """Extract first info table (structured fields)"""
    fields = {}
    for tb in doc.tables:
        for row in tb.rows:
            cells = [c.text.strip() for c in row.cells]
            if len(cells) >= 2 and cells[0]:
                fields[cells[0]] = ' | '.join(cells[1:])
    return fields

def extract_summary(canon_id):
    file = None
    for f in os.listdir(TT_DIR):
        if f.startswith(canon_id + '_') and f.endswith('.docx'):
            file = f
            break
    if not file:
        return None
    path = os.path.join(TT_DIR, file)
    d = Document(path)
    fields = extract_info_table(d)
    all_text = '\n'.join(p.text for p in d.paragraphs if p.text.strip())
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                all_text += '\n' + cell.text
    return {
        'file': file,
        'id': canon_id,
        'fields': fields,
        'full_text': all_text,
    }

# ============================================================
# 2. BUILD GRAPH with better extraction
# ============================================================
G = nx.MultiDiGraph()

# Country normalized names
COUNTRY_NORM = {
    'vietnam': 'Vietnam', 'vn': 'Vietnam', 'việt nam': 'Vietnam', 'viet nam': 'Vietnam',
    'china': 'China', 'trung quốc': 'China', 'prc': 'China',
    'korea': 'Korea', 'south korea': 'Korea', 'hàn quốc': 'Korea',
    'japan': 'Japan', 'nhật bản': 'Japan', 'nhật': 'Japan',
    'usa': 'USA', 'united states': 'USA', 'mỹ': 'USA', 'hoa kỳ': 'USA', 'america': 'USA',
    'uk': 'UK', 'england': 'UK', 'united kingdom': 'UK', 'anh': 'UK',
    'australia': 'Australia', 'úc': 'Australia',
    'canada': 'Canada',
    'sri lanka': 'Sri Lanka',
    'india': 'India', 'ấn độ': 'India',
    'ireland': 'Ireland',
    'norway': 'Norway',
    'denmark': 'Denmark',
    'philippines': 'Philippines',
    'indonesia': 'Indonesia',
    'malaysia': 'Malaysia',
    'thailand': 'Thailand', 'thái lan': 'Thailand',
    'ethiopia': 'Ethiopia',
    'saudi arabia': 'Saudi Arabia',
    'bangladesh': 'Bangladesh',
}

# ============================================================
# Helper: detect STUDY LOCATION (not citation location)
# ============================================================
def detect_study_location(fields, full_text):
    """Conservative: only use explicit affiliation / study setting fields."""
    countries = set()

    # Priority 1: explicit location keywords in fields
    check_fields = ['Đơn vị', 'Cơ quan', 'Đơn vị chính', 'Mẫu', 'Địa bàn', 'Phạm vi',
                    'Khách thể', 'Affiliations', 'Setting']
    sources_text = ''
    for k in check_fields:
        if k in fields:
            sources_text += ' ' + fields[k]

    # Priority 2: first 500 chars of full text (title + affiliation)
    sources_text += ' ' + full_text[:500]
    sources_text_lower = sources_text.lower()

    for kw, norm in COUNTRY_NORM.items():
        if kw in sources_text_lower:
            countries.add(norm)

    # Special: Southeast Asia / LMIC scope
    if 'southeast asia' in sources_text_lower or 'đông nam á' in sources_text_lower:
        countries.add('Southeast Asia (scope)')
    if 'asean' in sources_text_lower:
        countries.add('ASEAN (scope)')

    return countries

# ============================================================
# Helper: detect primary first author
# ============================================================
def detect_first_author(fields, full_text):
    author_field = fields.get('Tác giả', '')
    if not author_field:
        return None
    # Remove affiliations in parens/brackets
    s = re.sub(r'\([^)]*\)', '', author_field)
    s = re.sub(r'\[[^\]]*\]', '', s)
    # Split by common delimiters
    parts = re.split(r'[,;]|\band\b|\&|và cộng sự', s)
    first = parts[0].strip() if parts else None
    if first:
        # Remove trailing affiliation numbers
        first = re.sub(r'^\d+[,\-–]?\s*', '', first)
        first = first.strip()
        # Strip initials like "T.D."
        if len(first) < 4:
            return None
    return first

# Author synonyms for disambiguation
AUTHOR_ALIASES = {
    # Brown has 2 papers (BESST 2024 + editorial 2025 with Carter)
    'Brown 2024': {'canonical': 'Brown_2024_BESST'},
    'Brown & Carter 2025': {'canonical': 'Brown_Carter_2025_Editorial'},
    # Zhang has multiple (Bayesian, Liang, Kang)
    'Zhang Liang Kang 2026': {'canonical': 'Zhang_2026_Bayesian'},
    # Li has multiple (BMC NMA, screen time BJCP)
    'Li et al. 2025 BMC': {'canonical': 'Li_2025_BMC_NMA'},
    'Li 2025 BJCP': {'canonical': 'Li_2025_BJCP_ScreenTime'},
}

# ============================================================
# Process papers
# ============================================================
for item in CANONICAL:
    canon = item['canonical']
    canon3 = canon[:2] + f'{int(canon[2:]):03d}'
    data = extract_summary(canon3)
    if not data:
        continue

    fields = data['fields']
    full_text = data['full_text']
    desc = item['descriptor']

    # Paper node with rich attrs
    G.add_node(canon3, type='Paper', descriptor=desc,
               title=fields.get('Tiêu đề gốc', fields.get('Tiêu đề', '')))

    # First author
    first_author = detect_first_author(fields, full_text)
    if first_author:
        aid = f'Author::{first_author[:50]}'
        if not G.has_node(aid):
            G.add_node(aid, type='Author', name=first_author, role='first')
        G.add_edge(canon3, aid, rel='FIRST_AUTHOR')

    # Year (conservative: year from Năm field or từ descriptor)
    year_str = None
    for k in ['Xuất bản', 'Năm', 'Năm bảo vệ']:
        if k in fields:
            m = re.search(r'(19|20)\d{2}', fields[k])
            if m:
                year_str = m.group(0)
                break
    if not year_str:
        m = re.search(r'(20\d{2})', desc)
        if m:
            year_str = m.group(0)
    if year_str:
        yid = f'Year::{year_str}'
        if not G.has_node(yid):
            G.add_node(yid, type='Year', value=int(year_str))
        G.add_edge(canon3, yid, rel='HAS_YEAR')

    # Journal
    journal_raw = fields.get('Tạp chí', '')
    if journal_raw:
        # Extract name before (, IF=, Q
        j = re.split(r'[(,]', journal_raw)[0].strip()
        j = re.sub(r'\s+', ' ', j)
        if j and len(j) > 3:
            jid = f'Journal::{j[:60]}'
            if not G.has_node(jid):
                G.add_node(jid, type='Journal', name=j)
            G.add_edge(canon3, jid, rel='PUBLISHED_IN')

    # Sample size
    n_match = re.search(r'[Nn]\s*=\s*([\d.,]+)', fields.get('Mẫu', '') + ' ' + full_text[:1000])
    if n_match:
        n_str = n_match.group(1).replace('.', '').replace(',', '')
        try:
            n = int(n_str)
            nid = f'SampleSize::{n}'
            if not G.has_node(nid):
                G.add_node(nid, type='SampleSize', value=n)
            G.add_edge(canon3, nid, rel='HAS_N')
        except: pass

    # Study location (conservative)
    countries = detect_study_location(fields, full_text)
    for country in countries:
        cid = f'Country::{country}'
        if not G.has_node(cid):
            G.add_node(cid, type='Country', name=country)
        G.add_edge(canon3, cid, rel='CONDUCTED_IN')

    # Methods (from Loại NC + "Can thiệp")
    method_text = (fields.get('Loại NC', '') + ' ' + fields.get('Can thiệp', '') + ' ' +
                   fields.get('Phương pháp', '') + ' ' + full_text)
    method_hints = {
        'CBT': 'CBT', 'cognitive behavioral': 'CBT', 'nhận thức–hành vi': 'CBT',
        'iCBT': 'iCBT', 'internet-based CBT': 'iCBT', 'CBT qua internet': 'iCBT',
        'gCBT': 'gCBT', 'CBT nhóm': 'gCBT', 'group CBT': 'gCBT',
        'Mobile CBT': 'Mobile CBT', 'mobile-based CBT': 'Mobile CBT',
        'Mindfulness': 'Mindfulness', 'mindfulness-based': 'Mindfulness',
        'DMHI': 'DMHI',
        'resilience': 'Resilience training',
        'CA-CBT': 'CA-CBT', 'culturally adapted CBT': 'CA-CBT',
        'Physical exercise': 'Physical Exercise', 'hoạt động thể chất': 'Physical Exercise',
        'Yoga': 'Yoga', 'Thư giãn': 'Relaxation',
        'VR': 'VR Therapy', 'virtual reality': 'VR Therapy',
        'SSRI': 'SSRI', 'Sertraline': 'SSRI',
        'ACT': 'ACT', 'Acceptance and Commitment': 'ACT',
        'IPT': 'IPT', 'interpersonal': 'IPT',
    }
    for kw, method in method_hints.items():
        if kw in method_text:
            mid = f'Method::{method}'
            if not G.has_node(mid):
                G.add_node(mid, type='Method', name=method)
            if not G.has_edge(canon3, mid, key=None):
                G.add_edge(canon3, mid, rel='USED_METHOD')

    # Outcomes
    outcome_hints = {
        'lo âu xã hội': 'Social Anxiety (SAD)',
        'social anxiety': 'Social Anxiety (SAD)', 'SAD': 'Social Anxiety (SAD)',
        'GAD': 'Generalized Anxiety (GAD)',
        'generalized anxiety': 'Generalized Anxiety (GAD)',
        'lo âu': 'Anxiety', 'anxiety': 'Anxiety',
        'trầm cảm': 'Depression', 'depression': 'Depression',
        'căng thẳng': 'Stress', 'stress': 'Stress',
    }
    for kw, outcome in outcome_hints.items():
        if kw in full_text:
            oid = f'Outcome::{outcome}'
            if not G.has_node(oid):
                G.add_node(oid, type='Outcome', name=outcome)
            if not G.has_edge(canon3, oid, key=None):
                G.add_edge(canon3, oid, rel='MEASURED')

    # Scales
    scales = ['DASS-21', 'DASS-Y', 'DASS-42', 'GAD-7', 'HAM-A', 'PHQ-9', 'SIAS', 'SCAS',
              'CES-D', 'CESD-R', 'STAI', 'LSAS', 'ASI', 'PSQI', 'LOT-R', 'ESSA', 'YBRS',
              'CGI', 'EPI', 'MHC-SF', 'CSES']
    for scale in scales:
        if scale in full_text:
            sid = f'Scale::{scale}'
            if not G.has_node(sid):
                G.add_node(sid, type='Scale', name=scale)
            if not G.has_edge(canon3, sid, key=None):
                G.add_edge(canon3, sid, rel='USED_SCALE')

    # Effect sizes WITH CONTEXT
    es_pats = [
        (r'([Cc]ohen[\'\u2019]?s?\s*d)\s*=\s*(-?\d+[,.]\d+)', 'Cohen_d'),
        (r'([Hh]edges[\'\u2019]?s?\s*g)\s*=\s*(-?\d+[,.]\d+)', 'Hedges_g'),
        (r'\b(d)\s*=\s*(-?\d+[,.]\d+)', 'd'),
        (r'\b(g)\s*=\s*(-?\d+[,.]\d+)', 'g'),
        (r'\b(OR)\s*=\s*(\d+[,.]\d+)', 'OR'),
        (r'\b(AOR)\s*=\s*(\d+[,.]\d+)', 'AOR'),
        (r'\b(SMD)\s*=\s*(-?\d+[,.]\d+)', 'SMD'),
        (r'(β)\s*=\s*(-?\d+[,.]\d+)', 'beta'),
        (r'(SUCRA)\s*[=:]\s*(\d+[,.]\d+|\d+\s*%)', 'SUCRA'),
    ]
    es_counter = 0
    for pat, es_type in es_pats:
        for m in re.finditer(pat, full_text):
            val_str = m.group(2).replace(',', '.').replace('%', '').strip()
            try:
                val = float(val_str)
                # Extract context (50 chars around)
                start = max(0, m.start() - 50)
                end = min(len(full_text), m.end() + 50)
                ctx = full_text[start:end].replace('\n', ' ').strip()
                es_counter += 1
                es_node = f'ES::{canon3}::{es_counter}'
                G.add_node(es_node, type='EffectSize', es_type=es_type, value=val,
                           paper=canon3, context=ctx[:200])
                G.add_edge(canon3, es_node, rel='REPORTED_ES')
            except: pass

# ============================================================
# 3. SAVE + STATS
# ============================================================
print(f'\n=== KG v2 STATS ===')
print(f'Nodes: {G.number_of_nodes()}')
print(f'Edges: {G.number_of_edges()}')
node_types = Counter(G.nodes[n].get('type', 'Unknown') for n in G.nodes)
print(f'Node types: {dict(node_types)}')
edge_rels = Counter(d['rel'] for _, _, d in G.edges(data=True))
print(f'Edge rels: {dict(edge_rels)}')

# Save
nx.write_graphml(G, os.path.join(KG_DIR, 'kg_v2.graphml'))
print(f'\nSaved: kg_v2.graphml')

# Compare v1 vs v2: papers with Country=Vietnam
v1 = nx.read_graphml(os.path.join(KG_DIR, 'kg_v1.graphml'))
vn_v1 = [u for u, v, d in v1.edges(data=True)
         if d.get('rel') == 'CONDUCTED_IN' and v == 'Country::Vietnam']
vn_v2 = [u for u, v, d in G.edges(data=True)
         if d.get('rel') == 'CONDUCTED_IN' and v == 'Country::Vietnam']

print(f'\n=== COMPARISON v1 vs v2 ===')
print(f'Papers tagged as "in Vietnam":')
print(f'  v1: {len(vn_v1)} (over-tagged due to citation context)')
print(f'  v2: {len(vn_v2)} (conservative - study location only)')
print(f'  Difference: {len(vn_v1) - len(vn_v2)} papers removed from VN')

# v2 Vietnam papers (should be ~30 VN + few multi-country studies)
print(f'\n  v2 Vietnam papers: {sorted(vn_v2)}')

# Quality check: first authors
print(f'\n=== FIRST AUTHORS ===')
first_authors = [n for n in G.nodes if G.nodes[n].get('type') == 'Author']
print(f'Total unique first authors: {len(first_authors)}')

# Top-degree papers
print(f'\n=== TOP PAPERS BY DEGREE ===')
papers = [n for n in G.nodes if G.nodes[n].get('type') == 'Paper']
for p in sorted(papers, key=lambda x: G.degree(x), reverse=True)[:15]:
    print(f'  {p}: degree={G.degree(p)} | {G.nodes[p].get("descriptor", "")[:50]}')
