# -*- coding: utf-8 -*-
"""V3 — 2 doc dịch + phản biện chi tiết theo workflow chuẩn:
1. Herres & Ohannessian (2015) — Adolescent Coping Profiles
2. Steinhoff et al. (2023) — Longitudinal cohort 9 năm

Quy tắc (theo memory feedback_research_workflow):
- Tiếng Việt thuần, chú thích Anh trong ngoặc khi cần
- Đánh dấu trang (Trang X, Tạp chí, Vol)
- Bảng dữ liệu = bảng Word thật (tái tạo Table 1, 2, 3, 4, 5)
- Từ viết tắt: lần đầu giải thích inline + tô đỏ
- Reference list trích từ bài gốc (giữ tiếng Anh)
- Phản biện cuối bài: 4-5 điểm, mỗi điểm 3-5 câu ngắn, có dẫn chứng
- Phát hiện gap nghiên cứu chi tiết
- KHÔNG có phần "Truy vết" gây nhiễu
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien'

RED   = RGBColor(0xC0, 0, 0)
BLUE  = RGBColor(0, 0x70, 0xC0)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0, 0x70, 0x40)
ORANGE = RGBColor(0xE0, 0x6C, 0x00)  # cho đánh dấu trang
RED_BOLD = RGBColor(0xCC, 0x00, 0x00)  # cho từ viết tắt lần đầu

def make_doc():
    d = Document()
    s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
    s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
    for sec in d.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)
    return d

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(d, headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(9)
def title(d, text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(d, text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(d, text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color

def page_marker(d, text):
    """Đánh dấu trang — màu cam, in nghiêng, căn giữa."""
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(f'— — — {text} — — —')
    r.bold=True; r.italic=True; r.font.size=Pt(11)
    r.font.color.rgb=ORANGE; r.font.name='Times New Roman'

def acro(d, viet, eng, abbrev):
    """Từ viết tắt lần đầu — tô đỏ đậm."""
    p=d.add_paragraph()
    r=p.add_run(f'{viet} ({eng} — {abbrev})')
    r.bold=True; r.font.color.rgb=RED_BOLD; r.font.size=Pt(12); r.font.name='Times New Roman'

def crit_para(d, text):
    """Đoạn phản biện — tô đỏ."""
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run(text); r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'

FIGURES_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\02_Papers-goc\Coping_Effectiveness\figures'

def insert_image(d, fname, caption, width_cm=15):
    """Chèn hình từ figures/ + caption màu cam căn giữa."""
    import os as _os
    path = _os.path.join(FIGURES_DIR, fname)
    if not _os.path.exists(path):
        # Fallback: chỉ in cảnh báo nếu hình không tồn tại
        p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(f'[Hình không tải được: {fname}]')
        r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(11)
        return
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(); r.add_picture(path, width=Cm(width_cm))
    cap = d.add_paragraph(); cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rcap = cap.add_run(caption)
    rcap.italic = True; rcap.bold = True; rcap.font.size = Pt(10)
    rcap.font.color.rgb = ORANGE; rcap.font.name = 'Times New Roman'

# =====================================================================
# DOC 1 — HERRES & OHANNESSIAN 2015 — V3 CHI TIẾT
# =====================================================================
def make_doc_herres_v3():
    d = make_doc()

    # ========== TRANG BÌA ==========
    title(d, "BẢN DỊCH + PHẢN BIỆN CHI TIẾT (PHIÊN BẢN 3)", 14)
    title(d, "CÁC HỒ SƠ ỨNG PHÓ CỦA VỊ THÀNH NIÊN", 16)
    title(d, "PHÂN BIỆT BÁO CÁO TRIỆU CHỨNG TRẦM CẢM VÀ LO ÂU", 14)
    nr(d, "")
    subtitle(d, "Adolescent Coping Profiles Differentiate Reports of Depression and Anxiety Symptoms")
    nr(d, "")
    subtitle(d, "Joanna Herres & Christine McCauley Ohannessian (2015)")
    subtitle(d, "Journal of Affective Disorders, Tập 186: 312–319")
    subtitle(d, "DOI: 10.1016/j.jad.2015.07.031 — PMID 26275359 — PMC4565746")
    nr(d, "")
    subtitle(d, "Trợ lý nghiên cứu — 27/04/2026")
    subtitle(d, "Tiếng Việt thuần — chú thích thuật ngữ tiếng Anh trong ngoặc khi cần")
    subtitle(d, "Đường dẫn truy cập miễn phí: https://pmc.ncbi.nlm.nih.gov/articles/PMC4565746/")

    # ========== THÔNG TIN THƯ MỤC ==========
    H1(d, "THÔNG TIN THƯ MỤC")
    tbl(d, ['Mục', 'Nội dung'], [
        ['Tên bài (tiếng Anh)', 'Adolescent coping profiles differentiate reports of depression and anxiety symptoms'],
        ['Tạm dịch tên bài', 'Các hồ sơ ứng phó của vị thành niên phân biệt các báo cáo về triệu chứng trầm cảm và lo âu'],
        ['Tác giả', 'Joanna Herres + Christine McCauley Ohannessian'],
        ['Tạp chí', 'Journal of Affective Disorders (Elsevier)'],
        ['Tập / Số / Trang', 'Tập 186: trang 312–319 (số đặc biệt tháng 8/2015)'],
        ['Năm xuất bản', '2015'],
        ['DOI (mã định danh số)', '10.1016/j.jad.2015.07.031'],
        ['PMID (mã PubMed)', '26275359'],
        ['PMC (mã PubMed Central)', 'PMC4565746'],
        ['NIHMS (mã bản thảo NIH)', 'NIHMS714864'],
        ['Loại nghiên cứu', 'Nghiên cứu cắt ngang (cross-sectional) sử dụng phân tích hồ sơ tiềm ẩn (Latent Profile Analysis — LPA)'],
        ['Cỡ mẫu', 'N = 982 học sinh THPT (54% nữ; 46% nam)'],
        ['Tuổi (Mean ± SD)', '16,09 tuổi (SD = 0,68); lớp 10–11'],
        ['Bối cảnh',
         '7 trường THPT công lập vùng Mid-Atlantic Hoa Kỳ (Delaware, Pennsylvania, Maryland), khảo sát mùa xuân 2007'],
        ['Tài trợ', 'Viện Y tế Quốc gia Hoa Kỳ — National Institute of Health (mã grant K01-AA015059)'],
        ['Xung đột lợi ích', 'Không nêu trong bài'],
        ['PDF gốc trong dự án',
         '02_Papers-goc/Coping_Effectiveness/Herres_Ohannessian_2015_Adolescent_Coping_Profiles_JAD.pdf '
         '(791 KB, 21 trang, bản NIH manuscript NIHMS714864)'],
    ], [4.0, 12.0])

    # ========== BẢNG TỪ VIẾT TẮT ==========
    H1(d, "BẢNG TỪ VIẾT TẮT DÙNG TRONG TÀI LIỆU NÀY")
    nr(d, "Lần đầu mỗi từ viết tắt xuất hiện trong văn bản dưới đây sẽ được tô MÀU ĐỎ ĐẬM "
       "+ chú thích inline. Bảng tổng hợp:")
    tbl(d, ['Tiếng Việt', 'Tiếng Anh đầy đủ', 'Viết tắt'], [
        ['Phân tích hồ sơ tiềm ẩn', 'Latent Profile Analysis', 'LPA'],
        ['Liệu pháp nhận thức – hành vi', 'Cognitive Behavioural Therapy', 'CBT'],
        ['Bộ thang đo ứng phó (Carver 1989)', 'Coping Orientation to Problems Experienced — COPE Inventory', 'COPE'],
        ['Thang trầm cảm trẻ em (Trung tâm dịch tễ học)', 'Center for Epidemiological Studies Depression Scale for Children', 'CES-DC'],
        ['Thang sàng lọc lo âu trẻ em', 'Screen for Child Anxiety Related Disorders', 'SCARED'],
        ['Phân tích phương sai', 'Analysis of Variance', 'ANOVA'],
        ['Phân tích phương sai đa biến', 'Multivariate Analysis of Variance', 'MANOVA'],
        ['Tiêu chí thông tin Bayes', 'Bayesian Information Criterion', 'BIC'],
        ['Kiểm định tỉ số khả năng Vuong-Lo-Mendell-Rubin', 'Vuong-Lo-Mendell-Rubin Likelihood Ratio Test', 'VLMR'],
        ['Kiểm định Lo-Mendell-Rubin điều chỉnh', 'Adjusted Lo-Mendell-Rubin Test', 'Adj. LMR'],
        ['Hệ số tin cậy nội tại Cronbach', "Cronbach's alpha", 'α'],
        ['Hội đồng đạo đức nghiên cứu', 'Institutional Review Board', 'IRB'],
        ['Khoảng tin cậy 95%', '95% Confidence Interval', '95% CI'],
        ['Sai số chuẩn', 'Standard Deviation', 'SD'],
        ['Trung bình', 'Mean', 'M'],
        ['Số mẫu', 'Sample size', 'n hoặc N'],
        ['Tình trạng kinh tế xã hội', 'Socio-Economic Status', 'SES'],
    ], [5.0, 8.0, 3.0])

    # ========== PHẦN 1 — DẪN NHẬP ==========
    H1(d, "PHẦN 1 — DẪN NHẬP (INTRODUCTION)")
    page_marker(d, 'Trang 312, Journal of Affective Disorders, Tập 186, 8/2015')
    H2(d, "1.1. Bối cảnh lý thuyết")
    nr(d, "Tác giả mở đầu với định nghĩa kinh điển của Lazarus & Folkman (1984): ứng phó "
       "(coping) là cách thức cá nhân PHẢN ỨNG và QUẢN LÝ căng thẳng. Bài đặt nghiên cứu "
       "trong khung mô hình \"tố chất – căng thẳng\" (diathesis-stress model) — theo đó "
       "khác biệt cá nhân về ứng phó là yếu tố phân biệt giữa người phát triển triệu "
       "chứng trầm cảm sau căng thẳng và người vẫn giữ được khả năng phục hồi (Compas, "
       "Orosan & Grant 1993; Evans và cộng sự 2015).")
    nr(d, "Ứng phó được xem là YẾU TỐ BẢO VỆ — phát triển dần qua tuổi và được tinh chỉnh "
       "khi cá nhân học hỏi (Amirkhan & Auyeung 2007). Một số chiến lược tỏ ra hiệu quả "
       "hơn ở giai đoạn vị thành niên (Griffith, Dubow & Ippolito 2000), trong khi các "
       "chiến lược không thích nghi gắn liền với các vấn đề nội tâm hoá — gồm trầm cảm + "
       "lo âu (Compas và cộng sự 2001).")

    H2(d, "1.2. Phân loại ứng phó: Tiếp cận vs Tránh né")
    acro(d, "Phân tích hồ sơ tiềm ẩn", "Latent Profile Analysis", "LPA")
    nr(d, "Tác giả lập luận rằng thay vì xem xét từng chiến lược ứng phó riêng lẻ, NHẬN "
       "DIỆN CÁC LOẠI HÌNH (typologies) gồm các vị thành niên có MẪU HÌNH ứng phó tương "
       "tự nhau sẽ cho hiểu biết tốt hơn (Aldridge & Roesch 2008).")
    nr(d, "Hai nhóm chiến lược chính được phân biệt:")
    nr(d, "• ỨNG PHÓ TIẾP CẬN / GẮN KẾT (approach / engagement coping): bao gồm kiểm soát "
       "sơ cấp (giải quyết vấn đề) và kiểm soát thứ cấp (tái cấu trúc nhận thức), cùng "
       "với tìm kiếm hỗ trợ cảm xúc (Carver & Connor-Smith 2010; Skinner và cộng sự 2003).")
    nr(d, "• ỨNG PHÓ TRÁNH NÉ (avoidance coping) / TÁCH RỜI / TẬP TRUNG VÀO CẢM XÚC: liên "
       "quan đến phủ nhận hoặc lờ đi tác nhân căng thẳng — mang lại giảm nhẹ NGẮN HẠN "
       "nhưng làm tăng đau khổ DÀI HẠN (Lynch và cộng sự 2001).")

    H2(d, "1.3. Bằng chứng thực nghiệm trước đây")
    nr(d, "Tổng quan của Compas và cộng sự (2001) tìm thấy ứng phó tiếp cận liên quan với "
       "kết quả tích cực; ứng phó tránh né tiên lượng tăng trầm cảm và sử dụng chất gây "
       "nghiện (Cicognani 2011; Seiffge-Krenke 2000; Wills và cộng sự 2001). Lewis, Byrd "
       "& Ollendick (2012) phát hiện ứng phó tránh né ĐÓNG VAI TRÒ TRUNG GIAN trong mối "
       "liên hệ giữa căng thẳng và lo âu.")

    H2(d, "1.4. Vị thành niên là giai đoạn then chốt")
    nr(d, "Tuổi vị thành niên có những thách thức riêng: phát triển dậy thì, hình thành "
       "tự chủ, phát triển bản sắc, năng lực tự điều hoà tăng lên — cùng với các bối "
       "cảnh gia đình / bạn bè / trường học thay đổi liên tục (Smetana và cộng sự 2006). "
       "Tỉ lệ các vấn đề tâm lý tăng đáng kể trong giai đoạn này.")

    H2(d, "1.5. Khác biệt giới tính trong ứng phó")
    nr(d, "Các nghiên cứu trước cho thấy NỮ sử dụng phổ ứng phó RỘNG HƠN nam (Cicognani "
       "2011) và báo cáo dùng ứng phó nhiều hơn nói chung (Wilson, Pritchard & Revalee "
       "2005). Nữ thường dùng các chiến lược tiếp cận như TÌM HỖ TRỢ XÃ HỘI và GIẢI "
       "QUYẾT VẤN ĐỀ nhiều hơn (Eschenbeck, Kohlmann & Lohaus 2007; Seiffge-Krenke 2011), "
       "có thể phản ánh việc nữ coi trọng quan hệ liên cá nhân hơn (Rudolph 2002). "
       "Điều này có thể giải thích cho việc nữ có nguy cơ trầm cảm cao hơn nam (Kessler "
       "và cộng sự 2005).")

    H2(d, "1.6. Câu hỏi nghiên cứu")
    nr(d, "Tác giả đưa ra 3 câu hỏi:")
    nr(d, "(1) Bộ COPE có nhận diện được các hồ sơ ứng phó với mẫu hình ưu tiên chiến "
       "lược khác nhau ở vị thành niên hay không?")
    nr(d, "(2) Các hồ sơ này có khác nhau theo nhân khẩu học (giới tính, chủng tộc / sắc "
       "tộc, tuổi) không?")
    nr(d, "(3) Các loại hình hồ sơ có liên hệ KHÁC BIỆT với triệu chứng trầm cảm + lo "
       "âu không?")

    # ========== PHẦN 2 — PHƯƠNG PHÁP ==========
    H1(d, "PHẦN 2 — PHƯƠNG PHÁP (METHODS)")
    page_marker(d, 'Trang 313, Journal of Affective Disorders, Tập 186, 8/2015')

    H2(d, "2.1. Thiết kế nghiên cứu")
    nr(d, "Khảo sát cắt ngang (cross-sectional survey) tại 7 trường trung học phổ thông "
       "công lập vùng Mid-Atlantic Hoa Kỳ vào mùa xuân 2007.")

    H2(d, "2.2. Người tham gia")
    nr(d, "• Tổng cỡ mẫu: 982 học sinh THPT lớp 10-11")
    nr(d, "• Giới tính: 54% nữ, 46% nam")
    nr(d, "• Tuổi trung bình: 16,09 (SD = 0,68)")
    nr(d, "• Chủng tộc / sắc tộc: 65% Caucasian (da trắng), 19% gốc Phi (African-American), "
       "11% gốc Hispanic, 2% gốc Á (Asian), 3% khác")
    nr(d, "• Cấu trúc gia đình: 60% sống cùng cả 2 cha mẹ ruột; 91% sống với mẹ ruột; "
       "65% sống với cha ruột")
    nr(d, "• Trình độ phụ huynh: 96% mẹ + 97% cha tốt nghiệp THPT; 35% mẹ + 32% cha tốt "
       "nghiệp đại học; 9% mẹ + 8% cha học sau đại học / y khoa")

    H2(d, "2.3. Quy trình tuyển mẫu")
    nr(d, "• Xin đồng thuận chủ động từ phụ huynh (active parental consent — phụ huynh có "
       "thể từ chối: passive opt-out); chỉ 2 phụ huynh từ chối")
    nr(d, "• 45 vị thành niên từ chối tham gia")
    nr(d, "• Tỉ lệ tham gia tổng thể: 71% học sinh đủ điều kiện")
    nr(d, "• Hầu hết người không tham gia là do vắng mặt vào ngày khảo sát")
    nr(d, "• Hội đồng đạo đức nghiên cứu (Institutional Review Board — IRB) Đại học "
       "Delaware đã phê duyệt; có Giấy chứng nhận bảo mật của Chính phủ Hoa Kỳ "
       "(Certificate of Confidentiality)")
    nr(d, "• Khảo sát do nhân viên nghiên cứu được đào tạo + chứng nhận đạo đức nghiên "
       "cứu trên người thực hiện; thời gian ~40 phút; thưởng 1 vé xem phim")

    H2(d, "2.4. Thang đo")
    H3(d, "2.4.1. Thang đo ứng phó — Bộ COPE (Carver và cộng sự 1989)")
    acro(d, "Bộ thang đo ứng phó", "Coping Orientation to Problems Experienced — COPE Inventory", "COPE")
    nr(d, "60 mục, đo trên thang 1-4 (\"không làm gì cả\" → \"làm rất nhiều\"). 9 thang con "
       "được sử dụng — em liệt kê đầy đủ kèm hệ số tin cậy nội tại Cronbach alpha (α) và "
       "câu hỏi mẫu:")
    tbl(d, ['Thang con', 'Cronbach α', 'Câu hỏi mẫu (tiếng Anh)'], [
        ['Ứng phó chủ động (Active coping)', '0,75', '"I concentrate efforts on doing something about it"'],
        ['Phủ nhận (Denial)', '0,83', '"I say to myself this isn\'t real"'],
        ['Hỗ trợ xã hội cảm xúc (Emotional social support)', '0,86', '"I discuss feelings with someone"'],
        ['Hài hước (Humor)', '0,89', '— (không trích câu mẫu trong bài)'],
        ['Hỗ trợ xã hội công cụ (Instrumental social support)', '0,82', '"I try to get advice from someone"'],
        ['Tách rời tinh thần (Mental disengagement)', '0,55', '"I do other activities to take mind off things"'],
        ['Lập kế hoạch (Planning)', '0,82', '"I make a plan of action"'],
        ['Ứng phó tâm linh (Religious coping)', '0,90', '"I put my trust in God"'],
        ['Giải toả cảm xúc (Venting emotions)', '0,83', '"I get upset and let emotions out"'],
    ], [5.5, 1.5, 9.0])
    nr(d, "Lưu ý: thang con \"Tách rời tinh thần\" có Cronbach α = 0,55 (THẤP) — đây là hạn "
       "chế psychometric đáng chú ý.", italic=True, color=GRAY)

    H3(d, "2.4.2. Thang đo trầm cảm — CES-DC")
    acro(d, "Thang trầm cảm trẻ em", "Center for Epidemiological Studies Depression Scale for Children", "CES-DC")
    nr(d, "20 mục, đo trên thang 1-4 (\"không có\" → \"rất nhiều\"). Đánh giá cảm xúc / hành "
       "vi trong 1 TUẦN qua. Câu hỏi mẫu: \"I felt sad\". Tổng điểm khoảng 20-80. "
       "Cronbach α trong mẫu này = 0,91. Đã được kiểm định cho vị thành niên (Garrison "
       "và cộng sự 1991).")

    H3(d, "2.4.3. Thang đo lo âu — SCARED")
    acro(d, "Thang sàng lọc lo âu trẻ em", "Screen for Child Anxiety Related Disorders", "SCARED")
    nr(d, "41 mục, đo trên thang 0-2 (\"không đúng\" → \"rất đúng\"). Đánh giá triệu chứng "
       "trong 3 THÁNG qua. Câu hỏi mẫu: \"I get really frightened for no reason at all\". "
       "Tổng điểm khoảng 0-82. Cronbach α trong mẫu này = 0,94. Tính chất tâm trắc tốt "
       "(Birmaher và cộng sự 2003; Muris và cộng sự 2002).")

    H2(d, "2.5. Phân tích thống kê")
    page_marker(d, 'Trang 314, Journal of Affective Disorders, Tập 186, 8/2015')
    nr(d, "Phần mềm: Mplus phiên bản 6.0 (Muthén & Muthén 1998-2010).")
    nr(d, "Ước lượng: Maximum likelihood with robust standard errors (MLR) — phương pháp "
       "này robust với vi phạm giả định phân phối chuẩn (Yuan & Bentler 2000).")
    nr(d, "Quy trình LPA: Khớp các mô hình với 1, 2, 3, 4, 5 nhóm; so sánh chỉ số phù hợp "
       "để xác định số nhóm tối ưu.")
    nr(d, "Các chỉ số phù hợp được sử dụng:")
    acro(d, "Tiêu chí thông tin Bayes", "Bayesian Information Criterion", "BIC")
    nr(d, "+ BIC điều chỉnh theo cỡ mẫu — giá trị thấp hơn tốt hơn")
    acro(d, "Kiểm định tỉ số khả năng Vuong-Lo-Mendell-Rubin", "VLMR Likelihood Ratio Test", "VLMR")
    nr(d, "+ Kiểm định Lo-Mendell-Rubin điều chỉnh (Adj. LMR)")
    nr(d, "+ Entropy (thang 0-1; cao hơn = phân loại tốt hơn)")
    nr(d, "Phân tích so sánh: ANOVA cho khác biệt nhóm trên các thang COPE; kiểm định "
       "Tukey post-hoc cho so sánh từng cặp; kiểm định Chi-square cho khác biệt nhóm "
       "theo tuổi / chủng tộc / giới; MANOVA so sánh các nhóm về trầm cảm + lo âu, "
       "kiểm soát giới tính.")

    # ========== PHẦN 3 — KẾT QUẢ ==========
    H1(d, "PHẦN 3 — KẾT QUẢ (RESULTS)")
    page_marker(d, 'Trang 314-316, Journal of Affective Disorders, Tập 186, 8/2015')

    H2(d, "3.1. Thống kê mô tả các thang COPE (cả mẫu N=982)")
    tbl(d, ['Thang con COPE', 'Trung bình (M)', 'Sai số chuẩn (SD)'], [
        ['Ứng phó chủ động', '9,30', '2,85'],
        ['Phủ nhận', '6,53', '2,80'],
        ['Hỗ trợ xã hội cảm xúc', '9,72', '3,46'],
        ['Hài hước', '8,95', '3,50'],
        ['Hỗ trợ xã hội công cụ', '9,70', '3,24'],
        ['Tách rời tinh thần', '9,59', '2,67'],
        ['Lập kế hoạch', '9,48', '3,14'],
        ['Ứng phó tâm linh', '8,83', '3,96'],
        ['Giải toả cảm xúc', '8,73', '3,30'],
    ], [7.0, 4.0, 4.0])
    nr(d, "Tương quan giữa các thang: tất cả đều có ý nghĩa thống kê. Tương quan LỚN giữa "
       "các thang tiếp cận (vd: lập kế hoạch ↔ ứng phó chủ động: r = 0,78). Tương quan "
       "NHỎ giữa các thang tiếp cận và tránh né (vd: hỗ trợ cảm xúc ↔ phủ nhận: r = 0,29). "
       "Hỗ trợ cảm xúc ↔ hỗ trợ công cụ tương quan rất cao (r = 0,82, p < 0,05).")

    H2(d, "3.2. Khác biệt giới tính trên thang COPE")
    nr(d, "NỮ báo cáo cao hơn NAM có ý nghĩa thống kê trên 6 thang con. Em tóm tắt trong bảng:")
    tbl(d, ['Thang con', 'Nữ M (SD)', 'Nam M (SD)', 't', 'p'], [
        ['Hỗ trợ xã hội cảm xúc', '10,89 (3,27)', '8,36 (3,17)', 't(926) = 11,91', '< 0,01'],
        ['Hỗ trợ xã hội công cụ', '10,58 (3,10)', '8,69 (3,10)', 't(923) = 9,25', '< 0,01'],
        ['Ứng phó chủ động', '9,51 (2,82)', '9,06 (2,86)', 't(929) = 2,40', '< 0,05'],
        ['Tách rời tinh thần', '9,89 (2,59)', '9,25 (2,72)', 't(934) = 3,69', '< 0,01'],
        ['Ứng phó tâm linh', '9,23 (4,13)', '8,38 (3,69)', 't(917) = 3,26', '< 0,01'],
        ['Giải toả cảm xúc', '9,87 (3,20)', '7,40 (2,90)', 't(930) = 12,25', '< 0,01'],
    ], [4.5, 3.0, 3.0, 3.0, 1.5])
    nr(d, "Khác biệt LỚN NHẤT là giải toả cảm xúc (t=12,25) và hỗ trợ xã hội cảm xúc "
       "(t=11,91) — phù hợp với phát hiện kinh điển nữ giới sử dụng các chiến lược "
       "ứng phó dựa trên xã hội + cảm xúc nhiều hơn nam.")

    H2(d, "3.3. Kết quả phân tích hồ sơ tiềm ẩn (LPA) — Lựa chọn mô hình 4 nhóm")
    tbl(d, ['Số nhóm', 'BIC', 'Adj. BIC', 'VLMR p', 'Adj. LMR p', 'Entropy'], [
        ['1', '43.152,82', '43.095,65', '—', '—', '—'],
        ['2', '41.279,35', '41.190,42', '0,00', '0,00', '0,85'],
        ['3', '40.726,84', '40.606,15', '0,00', '0,00', '0,84'],
        ['4 (CHỌN)', '40.460,75', '40.308,30', '0,01', '0,01', '0,84'],
        ['5', '40.314,72', '40.130,52', '0,25', '0,25', '0,85'],
    ], [2.5, 2.5, 2.5, 2.5, 2.5, 2.5])
    nr(d, "QUYẾT ĐỊNH: chọn mô hình 4 nhóm. Lý do: Khi thử mô hình 5 nhóm, kiểm định tỉ "
       "số khả năng KHÔNG có ý nghĩa thống kê (VLMR p = 0,25; Adj. LMR p = 0,25) — nghĩa "
       "là thêm 1 nhóm KHÔNG cải thiện đáng kể độ phù hợp. Mô hình 4 nhóm là tối ưu.",
       bold=True)

    H2(d, "3.4. Bốn hồ sơ ứng phó — đặc điểm chi tiết (Bảng 4 tái tạo)")
    tbl(d, ['Thang con',
            'HS1: Đứt kết\n(n=153, 15,6%)',
            'HS2: Độc lập\n(n=375, 38,2%)',
            'HS3: Tìm hỗ trợ XH\n(n=300, 30,5%)',
            'HS4: Chủ động\n(n=122, 12,4%)'], [
        ['Ứng phó chủ động', '5,79 (1,64)', '8,72 (2,07)', '9,93 (1,99)', '13,19 (2,09)'],
        ['Phủ nhận', '5,07 (1,78)', '6,04 (2,18)', '6,79 (2,70)', '8,80 (3,81)'],
        ['Hỗ trợ XH cảm xúc', '4,91 (1,15)', '7,67 (1,61)', '11,70 (1,66)', '14,50 (1,60)'],
        ['Hài hước', '7,25 (3,44)', '8,54 (3,13)', '9,26 (3,33)', '11,23 (3,64)'],
        ['Hỗ trợ XH công cụ', '5,17 (1,35)', '7,96 (1,60)', '11,35 (1,57)', '14,42 (1,32)'],
        ['Tách rời tinh thần', '8,12 (2,83)', '8,94 (2,29)', '10,08 (2,34)', '11,53 (2,74)'],
        ['Lập kế hoạch', '5,56 (1,81)', '8,99 (2,38)', '10,13 (2,33)', '13,63 (2,06)'],
        ['Ứng phó tâm linh', '6,34 (2,88)', '8,33 (3,60)', '9,45 (3,85)', '11,33 (4,31)'],
        ['Giải toả cảm xúc', '5,93 (2,32)', '7,28 (2,49)', '9,79 (2,75)', '12,39 (2,79)'],
    ], [3.5, 3.0, 3.0, 3.0, 3.0])
    insert_image(d, 'Herres_2015_Figure1_COPE_profiles.png',
                 'Hình 1 (trang 15 PDF gốc) — Hồ sơ 4 nhóm ứng phó của vị thành niên trên 9 thang con của bộ COPE '
                 '(trích từ Herres & Ohannessian 2015, Journal of Affective Disorders, Tập 186)',
                 width_cm=15)
    nr(d, "Mô tả mỗi hồ sơ:", bold=True)
    nr(d, "• HS1 — ĐỨT KẾT (Disengaged Copers, n=153, 15,6%): mức ứng phó THẤP NHẤT "
       "tổng thể; có xu hướng dùng các chiến lược tránh né (hài hước, tách rời tinh "
       "thần) hơn các chiến lược tiếp cận.")
    nr(d, "• HS2 — ĐỘC LẬP (Independent Copers, n=375, 38,2%): mức ứng phó VỪA PHẢI; "
       "ít dùng giải toả cảm xúc và tìm hỗ trợ xã hội tương đối hơn.")
    nr(d, "• HS3 — TÌM HỖ TRỢ XÃ HỘI (Social Support Seeking Copers, n=300, 30,5%): "
       "ứng phó tổng thể CAO; ưu tiên rõ rệt các chiến lược tìm hỗ trợ xã hội.")
    nr(d, "• HS4 — CHỦ ĐỘNG (Active Copers, n=122, 12,4%): ứng phó CAO NHẤT trên TẤT "
       "CẢ các thang con; ưu tiên các chiến lược tiếp cận.")

    H2(d, "3.5. Khác biệt nhân khẩu học giữa các hồ sơ")
    nr(d, "• TUỔI: KHÔNG có khác biệt giữa các nhóm — Chi-square(6) = 6,07, p không có "
       "ý nghĩa thống kê.")
    nr(d, "• CHỦNG TỘC / SẮC TỘC: KHÔNG có khác biệt — Chi-square(12) = 16,10, p không "
       "có ý nghĩa thống kê.")
    nr(d, "• GIỚI TÍNH: CÓ khác biệt rõ rệt — Chi-square(3) = 89,30, p < 0,05.")
    tbl(d, ['Hồ sơ', '% Nam', '% Nữ'], [
        ['HS1 — Đứt kết', '66%', '34%'],
        ['HS2 — Độc lập', '60%', '40%'],
        ['HS3 — Tìm hỗ trợ XH', '34%', '66%'],
        ['HS4 — Chủ động', '25%', '75%'],
    ], [5.0, 3.0, 3.0])
    nr(d, "→ Nam tập trung nhiều ở 2 nhóm Đứt kết + Độc lập (ứng phó ÍT); Nữ tập "
       "trung nhiều ở 2 nhóm Tìm hỗ trợ XH + Chủ động (ứng phó NHIỀU).", bold=True)
    insert_image(d, 'Herres_2015_Figure2_COPE_by_gender.png',
                 'Hình 2 (trang 16 PDF gốc) — Điểm trung bình trên các thang COPE cho NỮ (trên) và NAM (dưới) '
                 'theo từng hồ sơ ứng phó (trích từ Herres & Ohannessian 2015)',
                 width_cm=15)

    H2(d, "3.6. Liên hệ giữa hồ sơ ứng phó và triệu chứng trầm cảm + lo âu (Bảng 5 tái tạo)")
    page_marker(d, 'Trang 316, Journal of Affective Disorders, Tập 186, 8/2015')
    nr(d, "Phân tích phương sai đa biến (MANOVA — Multivariate ANOVA) — kiểm soát giới "
       "tính. Các chỉ số dưới có chữ cái ghi chỉ số khác nhau (a, b, c) cho biết khác "
       "biệt có ý nghĩa thống kê (p < 0,05) giữa các nhóm:")
    tbl(d, ['Triệu chứng',
            'HS1: Đứt kết',
            'HS2: Độc lập',
            'HS3: Tìm hỗ trợ XH',
            'HS4: Chủ động',
            'F(3,785)', 'p'], [
        ['Trầm cảm tổng', '36,98a (12,39)', '33,69b (9,76)', '36,36a (10,90)', '37,42a (12,64)', '4,47', '0,01'],
        ['Lo âu xã hội', '3,77a (3,60)', '4,10a (3,30)', '4,64a (3,41)', '4,63a (3,66)', '2,64', '0,05'],
        ['Lo âu chia ly', '1,30a (2,11)', '1,63a (1,94)', '2,62b (2,45)', '3,24b (3,23)', '20,75', '0,01'],
        ['Lo âu lan toả', '3,88a (4,25)', '4,22a (3,66)', '5,76b (3,99)', '7,40c (4,92)', '20,82', '0,01'],
        ['LO ÂU TỔNG', '12,69a (12,45)', '13,29a (9,68)', '18,72b (12,13)', '22,44c (15,53)', '21,92', '0,01'],
    ], [3.0, 2.5, 2.5, 2.7, 2.5, 1.5, 1.0])

    nr(d, "PHÁT HIỆN CHÍNH:", bold=True, color=BLUE)
    nr(d, "(1) TRẦM CẢM: Nhóm ĐỘC LẬP (HS2) có trầm cảm THẤP NHẤT (M=33,69) — khác biệt "
       "có ý nghĩa thống kê so với cả 3 nhóm còn lại. 3 nhóm còn lại có trầm cảm tương "
       "đương nhau (~36-37 điểm).")
    nr(d, "(2) LO ÂU TỔNG: Nhóm CHỦ ĐỘNG (HS4) có lo âu CAO NHẤT (M=22,44) — gần GẤP "
       "ĐÔI nhóm Đứt kết (12,69) và Độc lập (13,29). Pattern này nhất quán cho lo âu "
       "chia ly và lo âu lan toả — KHÔNG cho lo âu xã hội (vốn có F nhỏ).")
    nr(d, "(3) GIỚI TÍNH KIỂM SOÁT: Tất cả khác biệt vẫn có ý nghĩa thống kê khi kiểm "
       "soát giới tính → khác biệt không phải do tỉ lệ giới khác nhau giữa các nhóm.")
    nr(d, "(4) HỆ SỐ R² điều chỉnh: dao động 0,006 (lo âu xã hội) đến 0,074 (lo âu tổng) "
       "— hồ sơ ứng phó CHỈ giải thích 0,6% đến 7,4% phương sai trong triệu chứng. Điều "
       "này cho thấy hồ sơ ứng phó là MỘT YẾU TỐ trong số nhiều yếu tố — không phải "
       "yếu tố duy nhất.", bold=True)

    # ========== PHẦN 4 — THẢO LUẬN ==========
    H1(d, "PHẦN 4 — THẢO LUẬN (DISCUSSION)")
    page_marker(d, 'Trang 316-318, Journal of Affective Disorders, Tập 186, 8/2015')

    H2(d, "4.1. Trả lời 3 câu hỏi nghiên cứu")
    nr(d, "Câu 1: LPA xác định được 4 nhóm vị thành niên khác biệt về MỨC ĐỘ ứng phó "
       "tổng thể. Trong từng nhóm, vị thành niên còn báo cáo SỞ THÍCH tương đối với "
       "một số chiến lược (vd: nhóm Đứt kết thiên về tránh né; nhóm Chủ động dùng "
       "tiếp cận nhiều hơn tránh né tương đối).")
    nr(d, "Câu 2: Các nhóm KHÔNG khác nhau về tuổi và chủng tộc. Khác biệt CÓ ý nghĩa "
       "về giới tính: nam tập trung ở 2 nhóm ứng phó ít (Đứt kết + Độc lập); nữ tập "
       "trung ở 2 nhóm ứng phó nhiều (Tìm hỗ trợ XH + Chủ động). Phù hợp với nghiên "
       "cứu trước (Seiffge-Krenke 2011).")
    nr(d, "Câu 3: Các hồ sơ KHÁC BIỆT có ý nghĩa về triệu chứng trầm cảm + lo âu. Phát "
       "hiện ĐÁNG CHÚ Ý: nhóm CHỦ ĐỘNG dù có ứng phó NHIỀU NHẤT lại báo cáo lo âu CAO "
       "NHẤT — phản trực giác.")

    H2(d, "4.2. Ba giả thuyết giải thích cho kết quả phản trực giác (Active Copers có lo âu cao)")
    nr(d, "Tác giả đưa ra 3 giả thuyết:")
    nr(d, "GIẢ THUYẾT 1 — TIẾP XÚC CĂNG THẲNG (Stress Exposure): Nhóm Chủ động có thể "
       "đang trải qua nhiều căng thẳng hơn nói chung → đó là lý do họ phải dùng nhiều "
       "ứng phó. Mức độ căng thẳng tổng thể có thể là biến thứ ba (third variable) "
       "giải thích cả ứng phó cao và lo âu cao.")
    nr(d, "GIẢ THUYẾT 2 — ỨNG PHÓ KHÔNG HIỆU QUẢ DO KIỂM SOÁT CHÚ Ý (Attentional "
       "Control): Nhóm Chủ động có thể có lo âu cao vì các nỗ lực ứng phó của họ không "
       "hiệu quả do khó khăn trong kiểm soát chú ý (Eysenck & Derakshan 2011).")
    nr(d, "GIẢ THUYẾT 3 — KHÔNG KHỚP TÁC NHÂN GÂY CĂNG THẲNG (Stressor Type Mismatch): "
       "Nhóm Chủ động có thể đang đối mặt với các tác nhân KHÔNG KIỂM SOÁT ĐƯỢC (vd: "
       "bệnh tật) — không phù hợp với chiến lược tiếp cận. Trong những tình huống "
       "không thể thay đổi, ứng phó tránh né có thể giảm lo âu tốt hơn (Dashora và "
       "cộng sự 2011).")

    H2(d, "4.3. Mẫu hình ứng phó lành mạnh — nhóm Độc lập")
    nr(d, "Đây là phát hiện đáng quan tâm: nhóm Độc lập có ÍT triệu chứng nhất dù ứng "
       "phó VỪA PHẢI (không phải nhiều nhất). Tác giả gợi ý quan hệ CONG (curvilinear) "
       "— người khoẻ mạnh dùng ứng phó vừa phải, người trầm cảm thường dùng ứng phó "
       "hoặc QUÁ NHIỀU hoặc QUÁ ÍT.")
    nr(d, "Hai cách diễn giải: (a) nhóm Độc lập trải qua ít căng thẳng hơn nên cần ít "
       "ứng phó hơn; (b) việc họ dùng các chiến lược tránh né nhiều hơn tương đối có "
       "thể KHỚP với loại tác nhân căng thẳng họ đối mặt (Griffith và cộng sự 2000).")

    H2(d, "4.4. Hệ quả lâm sàng")
    nr(d, "Tác giả nhấn mạnh: bác sĩ lâm sàng KHÔNG NÊN giả định vị thành niên có triệu "
       "chứng nội tâm hoá cao là KHÔNG có ứng phó. Trong thực tế, họ có thể đang dùng "
       "ứng phó (kể cả thích nghi) nhưng KHÔNG đủ. Thay vì dạy kỹ năng ứng phó CHUNG "
       "CHUNG, can thiệp nên giúp vị thành niên CHỌN ĐÚNG TỔ HỢP chiến lược cho từng "
       "loại tác nhân căng thẳng cụ thể.", bold=True)

    # ========== PHẦN 5 — DANH MỤC TÀI LIỆU THAM KHẢO ==========
    H1(d, "PHẦN 5 — DANH MỤC TÀI LIỆU THAM KHẢO TRONG BÀI GỐC")
    nr(d, "Bài gốc Herres & Ohannessian (2015) có TỔNG CỘNG 63 tài liệu tham khảo. Em "
       "liệt kê 20 tài liệu đầu — phần em đã đối chiếu trực tiếp từ PMC NCBI.",
       italic=True, color=GRAY)
    nr(d, "(Giữ nguyên tiếng Anh — định dạng gốc của tạp chí Journal of Affective Disorders)",
       italic=True, color=GRAY)

    refs = [
        "Aldridge AA, Roesch SC. Developing coping typologies of minority adolescents: A latent profile analysis. Journal of Adolescence. 2008;31(4):499–517.",
        "Amirkhan J, Auyeung B. Coping with stress across the lifespan: Absolute vs. relative changes in strategy. Journal of Applied Developmental Psychology. 2007;28(4):298–317.",
        "Barber BK, Olsen JA. Assessing the transitions to middle and high school. Journal of Adolescent Research. 2004;19(1):3–30.",
        "Bekker MH, van Mens-Verhulst J. Anxiety disorders: Sex differences in prevalence, degree, and background, but gender-neutral treatment. Gender Medicine. 2007;4:S178–S193.",
        "Birmaher B, Khetarpal S, Cully M, Brent DA, McKenzie S. Screen for child anxiety related disorders (SCARED). In: VandeCreek L, Jackson TL, eds. Innovations in Clinical Practice: Focus on Children & Adolescents. Sarasota, FL: Professional Resource Press; 2003:99–104.",
        "Carver CS, Connor-Smith J. Personality and coping. Annual Review of Psychology. 2010;61:679–704.",
        "Carver CS, Scheier MF, Weintraub JK. Assessing coping strategies: A theoretically based approach. Journal of Personality and Social Psychology. 1989;56:267–283.",
        "Cicognani E. Coping strategies with minor stressors in adolescence: Relationships with social support, self-efficacy, and psychological well-being. Journal of Applied Social Psychology. 2011;41(3):559–578.",
        "Clarke AT. Coping with interpersonal stress and psychosocial health among children and adolescents: A meta-analysis. Journal of Youth and Adolescence. 2006;35(1):11–24.",
        "Compas BE, Orosan PG, Grant KE. Adolescent stress and coping: Implications for psychopathology during adolescence. Journal of Adolescence. 1993;16(3):331–349.",
        "Compas BE, Connor-Smith J, Saltzman H, Thomsen AH, Wadsworth ME. Coping with stress during childhood and adolescence: Problems, progress, and potential in theory and research. Psychological Bulletin. 2001;127(1):87–127.",
        "Connor-Smith JK, Compas BE, Wadsworth ME, Thomsen AH, Saltzman H. Responses to stress in adolescence: Measurement of coping and involuntary stress responses. Journal of Consulting and Clinical Psychology. 2000;68(6):976.",
        "Dashora P, Erdem G, Slesnick N. Better to bend than to break: Coping strategies utilized by substance-abusing homeless youth. Journal of Health Psychology. 2011;16(1):158–168.",
        "Dekovic M, Ten Have M, Vollebergh W, et al. The cross-cultural equivalence of parental rearing measure: EMBU-C. European Journal of Psychological Assessment. 2006;22:85–91.",
        "Duncan DF. Growing up under the gun: Children and adolescents coping with violent neighborhoods. Journal of Primary Prevention. 1996;16:343–356.",
        "Essau CA, Conradt J, Sasagawa S, Ollendick TH. Prevention of anxiety symptoms in children: Results from a universal school-based trial. Behavior Therapy. 2012;43(2):450–464.",
        "Eschenbeck H, Kohlmann CW, Lohaus A. Gender differences in coping strategies in children and adolescents. Journal of Individual Differences. 2007;28(1):18.",
        "Evans LD, Kouros C, Frankel SA, et al. Longitudinal relations between stress and depressive symptoms in youth: Coping as a mediator. Journal of Abnormal Child Psychology. 2015;43:355–368.",
        "Eysenck MW, Derakshan N. New perspectives in attentional control theory. Personality and Individual Differences. 2011;50(7):955–960.",
        "Gamble WC. Perceptions of controllability and other stressor event characteristics as determinants of coping among young adolescents and young adults. Journal of Youth and Adolescence. 1994;23(1):65–84.",
    ]
    for i, r in enumerate(refs, 1):
        nr(d, f"{i}. {r}", size=11)
    nr(d, "(Còn 43 tài liệu tham khảo tiếp theo trong bài gốc — thầy có thể truy cập đầy "
       "đủ tại PMC4565746)", italic=True, color=GRAY, size=10)

    # ========== PHẦN 6 — PHẢN BIỆN CHI TIẾT ==========
    H1(d, "PHẦN 6 — PHẢN BIỆN CHI TIẾT (CRITICAL REVIEW)")
    nr(d, "Theo workflow chuẩn: 5 điểm chính, mỗi điểm có dẫn chứng cụ thể.",
       italic=True, color=GRAY)

    H2(d, "Điểm 1 — Phương pháp LPA là điểm mạnh chính của bài")
    crit_para(d, "Bài sử dụng phân tích hồ sơ tiềm ẩn (LPA) — tiếp cận \"tập trung vào "
              "cá nhân\" (person-centered). Điều này khác với hầu hết nghiên cứu trước "
              "đó tập trung vào biến đơn lẻ. LPA cho phép nắm bắt sự đa dạng \"phong cách\" "
              "ứng phó tự nhiên trong dân số. Cỡ mẫu N=982 đủ lớn cho LPA ổn định "
              "(Aldridge & Roesch 2008). Wen et al. 2020 (corpus dự án QT08) cũng dùng "
              "LPA trên 900 học sinh Trung Quốc nông thôn — cùng triết lý phương pháp.")

    H2(d, "Điểm 2 — Phát hiện phản trực giác về \"ứng phó nhiều = lo âu cao\" cần thận trọng")
    crit_para(d, "Phát hiện nhóm Chủ động (ứng phó NHIỀU NHẤT) có lo âu CAO NHẤT (M=22,44 — "
              "gần gấp đôi nhóm Đứt kết 12,69 và Độc lập 13,29) là phản trực giác. Tuy "
              "nhiên thiết kế cắt ngang KHÔNG cho phép kết luận nhân quả. Có 3 khả năng "
              "khác nhau — em đã tổng hợp ở mục 4.2. Điều quan trọng: KHÔNG nên kết "
              "luận \"dạy ứng phó nhiều là có hại\" từ bài này. Cần nghiên cứu DÀI HẠN "
              "(longitudinal) như Steinhoff 2023 (đã dịch riêng) để biết hướng nhân quả.")

    H2(d, "Điểm 3 — Hạn chế thiết kế cắt ngang (cross-sectional) — không thể kết luận nhân quả")
    crit_para(d, "Tác giả tự thừa nhận hạn chế này trong phần Limitations. Có 3 cách giải "
              "thích quan hệ ứng phó ↔ lo âu mà thiết kế cắt ngang không phân biệt được: "
              "(a) ứng phó ảnh hưởng đến lo âu — giả thuyết tác giả; (b) lo âu thúc đẩy "
              "người ta dùng nhiều ứng phó — quan hệ nhân quả ngược (reverse causation); "
              "(c) cả hai cùng ảnh hưởng lẫn nhau (bidirectional). Steinhoff 2023 longitudinal "
              "9 năm (đã dịch riêng) cho thấy tác động NHỎ — phù hợp khả năng (a) yếu hơn (b).")

    H2(d, "Điểm 4 — Cronbach alpha thang \"Tách rời tinh thần\" RẤT THẤP — vấn đề psychometric")
    crit_para(d, "Trong 9 thang con COPE, thang \"Tách rời tinh thần\" (Mental disengagement) "
              "có α = 0,55 — DƯỚI ngưỡng chấp nhận thông thường 0,70. Điều này có nghĩa "
              "thang đo có vấn đề về tin cậy nội tại. Ảnh hưởng đến độ tin cậy của LPA "
              "vì LPA dựa trên 9 thang con — 1 thang yếu có thể ảnh hưởng đến phân loại "
              "nhóm. Tác giả KHÔNG thảo luận chi tiết hạn chế này trong bài.")

    H2(d, "Điểm 5 — Mẫu Mid-Atlantic Hoa Kỳ — generalisability hạn chế cho Việt Nam")
    crit_para(d, "Mẫu thuộc 7 trường công lập Delaware, Pennsylvania, Maryland — vùng "
              "Mid-Atlantic Hoa Kỳ. 65% Caucasian. Văn hoá / hệ thống giáo dục / cơ chế "
              "hỗ trợ tâm lý hoàn toàn khác Việt Nam. Phát hiện \"hỗ trợ xã hội nhiều = "
              "lo âu cao\" có thể KHÁC ở Việt Nam — văn hoá Á Đông coi trọng gia đình "
              "+ tập thể, hỗ trợ xã hội có thể là yếu tố BẢO VỆ chứ không phải chỉ báo "
              "căng thẳng. Cần nghiên cứu nhân rộng (replication) ở Việt Nam.")

    # ========== PHẦN 7 — PHÁT HIỆN GAP ==========
    H1(d, "PHẦN 7 — PHÁT HIỆN GAP NGHIÊN CỨU TỪ BÀI NÀY")

    H2(d, "Gap 1 — Thiếu nghiên cứu LPA về ứng phó ở học sinh Việt Nam")
    nr(d, "Bài Herres 2015 là ở Hoa Kỳ. Wen 2020 (QT08 corpus dự án) là ở Trung Quốc "
       "nông thôn. Việt Nam CHƯA CÓ nghiên cứu LPA nào về ứng phó của học sinh trung "
       "học. Có thể là chủ đề nghiên cứu sinh / luận văn cao học rất phù hợp.")

    H2(d, "Gap 2 — Thiếu thang đo COPE bản tiếng Việt được kiểm định")
    nr(d, "Bộ COPE 60 mục của Carver 1989 chưa có phiên bản Việt được kiểm định chính "
       "thức. Cần dịch + back-translate + kiểm định độ tin cậy nội tại + cấu trúc nhân "
       "tố trên mẫu Việt Nam. Đây là bước CẦN trước khi thực hiện LPA ở Việt Nam.")

    H2(d, "Gap 3 — Thiếu nghiên cứu kết hợp ứng phó + ĐO MỨC ĐỘ CĂNG THẲNG đồng thời")
    nr(d, "Hạn chế chính của Herres 2015: KHÔNG đo mức độ căng thẳng. Vì vậy không phân "
       "biệt được \"ứng phó nhiều vì stress nhiều\" hay \"ứng phó nhiều gây ra anxiety\". "
       "Nghiên cứu Việt Nam tương lai cần đo CẢ căng thẳng (vd: thang Adolescent Stress "
       "Questionnaire — ASQ; Byrne et al. 2007) đồng thời với ứng phó.")

    H2(d, "Gap 4 — Thiếu nghiên cứu DÀI HẠN ở vị thành niên Việt Nam")
    nr(d, "Cohort dài hạn z-proso ở Thuỵ Sĩ (Steinhoff 2023) theo dõi 9 năm. Việt Nam "
       "có cohort thanh thiếu niên dài hạn nào không? Nếu chưa, cần khởi động — đặc biệt "
       "có giá trị để nghiên cứu các giai đoạn chuyển tiếp (lớp 9 → lớp 10, lớp 12 → "
       "đại học, đại học → việc làm).")

    H2(d, "Gap 5 — Thiếu nghiên cứu liên hệ HỒ SƠ ỨNG PHÓ với CAN THIỆP")
    nr(d, "Nếu xác định được 4 hồ sơ ở Việt Nam, câu hỏi tiếp theo: nên thiết kế can "
       "thiệp NÀO cho NHÓM NÀO? Vd: nhóm Đứt kết cần kích hoạt hành vi (behavioral "
       "activation); nhóm Chủ động cần được dạy CHỌN chiến lược phù hợp; nhóm Tìm hỗ "
       "trợ XH cần kỹ năng giải quyết vấn đề độc lập. Có thể là nghiên cứu can thiệp "
       "phân tầng theo hồ sơ (profile-tailored intervention) — chưa ai làm ở Việt Nam.")

    H2(d, "Gap 6 — Thiếu so sánh hồ sơ ứng phó GIỮA các nhóm tuổi")
    nr(d, "Bài Herres 2015 chỉ HS 16 tuổi. Hồ sơ ứng phó có khác nhau ở lớp 6 (12 tuổi), "
       "lớp 9 (15 tuổi), lớp 12 (18 tuổi)? Có theo trajectory phát triển không? Cần "
       "nghiên cứu cross-sectional + longitudinal để vẽ trajectories — đặc biệt ở thời "
       "điểm chuyển cấp THCS-THPT có nhiều stress mới.")

    out = os.path.join(OUT_DIR, 'Herres_Ohannessian_2015_CopingProfiles_v3_chi_tiet_27042026.docx')
    d.save(out); return out

# =====================================================================
# DOC 2 — STEINHOFF 2023 — V3 CHI TIẾT
# =====================================================================
def make_doc_steinhoff_v3():
    d = make_doc()

    # ========== TRANG BÌA ==========
    title(d, "BẢN DỊCH + PHẢN BIỆN CHI TIẾT (PHIÊN BẢN 3)", 13)
    title(d, "CÁC YẾU TỐ TIÊN LƯỢNG TỪ TUỔI VỊ THÀNH NIÊN SỚM", 14)
    title(d, "ĐẾN ỨNG PHÓ THÍCH NGHI VÀ CĂNG THẲNG TÂM LÝ", 14)
    title(d, "Ở THANH NIÊN TRONG ĐẠI DỊCH COVID-19", 14)
    title(d, "(NGHIÊN CỨU COHORT DÀI HẠN 9 NĂM)", 13)
    nr(d, "")
    subtitle(d, "Early Adolescent Predictors of Young Adults' Distress and Adaptive Coping")
    subtitle(d, "During the COVID-19 Pandemic: Findings From a Longitudinal Cohort Study")
    nr(d, "")
    subtitle(d, "Steinhoff A, Johnson-Ferguson L, Bechtiger L và cộng sự (2023)")
    subtitle(d, "Journal of Early Adolescence, Tập 44(9): 1250–1280")
    subtitle(d, "DOI: 10.1177/02724316231181660 — PMID 39372429 — PMC10261967")
    nr(d, "")
    subtitle(d, "Trợ lý nghiên cứu — 27/04/2026 — V3 chi tiết")
    subtitle(d, "Tiếng Việt thuần — chú thích thuật ngữ tiếng Anh trong ngoặc khi cần")
    subtitle(d, "Đường dẫn truy cập miễn phí: https://pmc.ncbi.nlm.nih.gov/articles/PMC10261967/")

    # ========== THÔNG TIN THƯ MỤC ==========
    H1(d, "THÔNG TIN THƯ MỤC")
    tbl(d, ['Mục', 'Nội dung'], [
        ['Tên bài (tiếng Anh)',
         "Early Adolescent Predictors of Young Adults' Distress and Adaptive Coping During "
         "the COVID-19 Pandemic: Findings From a Longitudinal Cohort Study"],
        ['Tạm dịch tên bài',
         'Các yếu tố tiên lượng từ tuổi vị thành niên sớm cho căng thẳng tâm lý và ứng phó '
         'thích nghi của thanh niên trong đại dịch COVID-19: Phát hiện từ một nghiên cứu '
         'cohort dài hạn'],
        ['Tác giả', 'Steinhoff A, Johnson-Ferguson L, Bechtiger L, Murray AL, Hepp U, '
         'Ribeaud D, Eisner M, Shanahan L'],
        ['Tạp chí', 'Journal of Early Adolescence (Sage)'],
        ['Tập / Số / Trang', 'Tập 44(9): trang 1250–1280'],
        ['Năm xuất bản (online)', '2023 (xuất bản online tháng 6/2023)'],
        ['DOI', '10.1177/02724316231181660'],
        ['PMID', '39372429'],
        ['PMC', 'PMC10261967'],
        ['Loại nghiên cứu', 'Nghiên cứu cohort dài hạn (longitudinal cohort study) theo dõi '
         '9 năm — từ tuổi 13 (2011) đến tuổi ~22 (4-9/2020 trong đại dịch COVID-19)'],
        ['Cỡ mẫu phân tích', 'N = 786 (mẫu gốc cohort z-proso ban đầu là 1.675 học sinh '
         'lớp 1 của Zurich năm 2004)'],
        ['Bối cảnh', 'Zurich, Thuỵ Sĩ. Dự án Phát triển Xã hội Zurich (Zurich Project on '
         'Social Development from Childhood to Adulthood — z-proso)'],
        ['Tài trợ', 'Quỹ Jacobs (mã 2010-888, 2013-1081-1) + Quỹ Khoa học Quốc gia Thuỵ Sĩ '
         '(SNSF mã 100014_132124, 100014_149979, 10FI14_170402/2, 10FI14_170409)'],
        ['Xung đột lợi ích', 'Tác giả công bố KHÔNG có xung đột lợi ích'],
        ['PDF gốc trong dự án',
         '02_Papers-goc/Coping_Effectiveness/Steinhoff_2023_Longitudinal_Coping_COVID_JEA.pdf '
         '(1,1 MB, 31 trang)'],
    ], [4.0, 12.0])

    # ========== BẢNG TỪ VIẾT TẮT ==========
    H1(d, "BẢNG TỪ VIẾT TẮT DÙNG TRONG TÀI LIỆU NÀY")
    tbl(d, ['Tiếng Việt', 'Tiếng Anh đầy đủ', 'Viết tắt'], [
        ['Dự án Phát triển Xã hội Zurich', 'Zurich Project on Social Development', 'z-proso'],
        ['Bộ thang đo COPE rút gọn', 'Brief COPE Inventory', 'Brief COPE'],
        ['Hệ thống chỉ số kinh tế xã hội quốc tế', 'International Socio-Economic Index', 'ISEI'],
        ['Hệ số tin cậy nội tại Cronbach', "Cronbach's alpha", 'α'],
        ['Hệ số chuẩn hoá', 'Standardized regression coefficient', 'β'],
        ['Khoảng tin cậy 95%', '95% Confidence Interval', '95% CI'],
        ['Hệ số tương quan', 'Correlation coefficient', 'r'],
        ['Trung bình', 'Mean', 'M'],
        ['Sai số chuẩn', 'Standard Deviation', 'SD'],
        ['Phép ước lượng khả năng cực đại với sai số chuẩn vững',
         'Maximum Likelihood with Robust standard errors', 'MLR'],
        ['Tình trạng kinh tế xã hội', 'Socio-Economic Status', 'SES'],
        ['Phép ngẫu nhiên hoàn toàn của dữ liệu thiếu', 'Missing Completely At Random', 'MCAR'],
        ['Multiple Imputation', 'Multiple Imputation', 'MI'],
        ['Bảng câu hỏi nuôi dạy con (Alabama)', 'Alabama Parenting Questionnaire', 'APQ'],
        ['Bảng câu hỏi hành vi xã hội (Tremblay)', 'Social Behavior Questionnaire', 'SBQ'],
    ], [5.5, 7.5, 3.0])

    # ========== PHẦN 1 — DẪN NHẬP ==========
    H1(d, "PHẦN 1 — DẪN NHẬP (INTRODUCTION)")
    page_marker(d, 'Trang 1250-1253, Journal of Early Adolescence, Tập 44(9), 2023')

    H2(d, "1.1. Bối cảnh khủng hoảng SKTT thanh thiếu niên")
    nr(d, "Tác giả mở đầu nhấn mạnh phản ứng của người trẻ với các sự kiện căng thẳng "
       "(quá trình cảm xúc + đánh giá nhận thức) và các phản ứng ứng phó thích nghi "
       "đóng vai trò QUAN TRỌNG trong nguy cơ + khả năng phục hồi sức khỏe tâm thần "
       "(Horwitz và cộng sự 2011; Vannucci và cộng sự 2018).")
    nr(d, "Các chỉ báo căng thẳng tâm lý — sự suy giảm sức khỏe tâm thần (well-being), "
       "tuyệt vọng (hopelessness), cảm nhận xáo trộn cuộc sống do tác nhân mới như "
       "COVID-19 — báo hiệu nguy cơ SKTT tăng. Ngược lại, các chiến lược ứng phó "
       "thích nghi giúp DUY TRÌ hoặc CẢI THIỆN sức khỏe tâm thần dưới căng thẳng.")
    nr(d, "Tác giả trích lời cảnh báo về \"khủng hoảng SKTT vị thành niên + thanh "
       "niên chưa được đảo ngược\" (Gunnell và cộng sự 2018; Báo cáo Tổng Y sĩ Hoa "
       "Kỳ 2021). Hiểu các tiền đề phát triển của phản ứng căng thẳng và ứng phó là "
       "thiết yếu để can thiệp.")

    H2(d, "1.2. Bối cảnh COVID-19 ở Thuỵ Sĩ")
    nr(d, "Trong đợt phong toả tháng 3/2020 ở Thuỵ Sĩ, các hạn chế bao gồm: cấm tụ "
       "tập, đóng cửa đại học + cửa hàng, đóng biên giới, yêu cầu làm việc tại nhà "
       "(Kohler và cộng sự 2020). Các hạn chế này XUNG ĐỘT với nhu cầu phát triển "
       "của thanh niên: kết nối bạn bè / người yêu, tiến bộ trong giáo dục + sự nghiệp.")

    H2(d, "1.3. Khái niệm ứng phó")
    acro(d, "Bộ thang đo COPE rút gọn", "Brief COPE Inventory", "Brief COPE")
    nr(d, "Theo Compas và cộng sự (2017) và Folkman (2012), ứng phó là các chiến lược "
       "NHẬN THỨC và HÀNH VI để quản lý các yêu cầu do tác nhân căng thẳng đặt ra — "
       "bao gồm nhưng không giới hạn ở việc điều hoà cảm xúc tiêu cực. Chiến lược "
       "khác nhau theo NGUỒN (tự túc vs có hỗ trợ xã hội) và TRỌNG TÂM (điều hoà cảm "
       "xúc, tạo nghĩa, giải quyết vấn đề).")

    H2(d, "1.4. Hai giả thuyết cạnh tranh về sự kiện căng thẳng sớm")
    nr(d, "Tác giả đặt vấn đề về 2 giả thuyết cạnh tranh về tác động của các sự kiện "
       "căng thẳng đầu đời lên sức đề kháng sau này:")
    nr(d, "GIẢ THUYẾT NHẠY CẢM (Sensitization Hypothesis — Rutter 2012): Trải nghiệm "
       "căng thẳng tích luỹ ĐẦU ĐỜI làm TĂNG nhạy cảm sinh lý với căng thẳng → tăng "
       "tổn thương khi gặp tác nhân căng thẳng MỚI → nội tâm hoá xấu đi (Low và cộng "
       "sự 2012).")
    nr(d, "GIẢ THUYẾT MIỄN DỊCH TÂM LÝ (Inoculation Hypothesis — Rutter 2012): Việc "
       "ĐỐI MẶT THÀNH CÔNG với căng thẳng đầu đời có thể GIẢM nhạy cảm với tác nhân "
       "căng thẳng tương lai và MỞ RỘNG kho ứng phó thích nghi (Oldehinkel và cộng "
       "sự 2014).")
    nr(d, "Đây là một trong những đóng góp lý thuyết QUAN TRỌNG của bài — kiểm định "
       "trực tiếp 2 giả thuyết đối lập trên cùng một cohort dài hạn.", bold=True)

    H2(d, "1.5. Vai trò tương tác hỗ trợ cha-mẹ-con")
    nr(d, "Tương tác hỗ trợ giữa cha mẹ và con được biết là yếu tố \"resilience\" "
       "(khả năng phục hồi). Cha mẹ phản ứng hỗ trợ với cảm xúc tiêu cực giúp con "
       "phát triển điều hoà cảm xúc — nền tảng của ứng phó với tác nhân căng thẳng "
       "mới (Perry và cộng sự 2020; Spinrad và cộng sự 2020). Cha mẹ vừa khuyến khích "
       "rõ ràng (explicit) vừa làm gương ngầm (implicit) cho các chiến lược ứng phó "
       "(Tu và cộng sự 2020). Gia đình hỗ trợ dạy con rằng hỗ trợ xã hội luôn sẵn "
       "có — kỹ năng này chuyển sang quan hệ bạn bè / người yêu sau này (Skinner & "
       "Zimmer-Gembeck 2007).")

    H2(d, "1.6. Mục tiêu nghiên cứu")
    nr(d, "Xem xét liệu các yếu tố ở tuổi vị thành niên sớm (13 tuổi) — bao gồm:")
    nr(d, "• Triệu chứng nội tâm hoá (anxiety + depression symptoms)")
    nr(d, "• Sự kiện căng thẳng tích luỹ (cumulative stressful life events)")
    nr(d, "• Tương tác hỗ trợ cha-mẹ-con (supportive parent-child interactions)")
    nr(d, "có TIÊN LƯỢNG được căng thẳng tâm lý + ứng phó thích nghi của thanh niên "
       "(22 tuổi) trong giai đoạn phong toả COVID-19 đầu tiên ở Thuỵ Sĩ hay không.")
    insert_image(d, 'Steinhoff_2023_Figure1_conceptual_framework.png',
                 'Hình 1 (trang 6 PDF gốc) — Khung khái niệm: các yếu tố tiên lượng từ tuổi vị thành niên sớm cho '
                 'căng thẳng tâm lý và ứng phó thích nghi của thanh niên (trích từ Steinhoff và cộng sự 2023, '
                 'Journal of Early Adolescence, Tập 44(9))',
                 width_cm=15)

    # ========== PHẦN 2 — PHƯƠNG PHÁP ==========
    H1(d, "PHẦN 2 — PHƯƠNG PHÁP (METHODS)")
    page_marker(d, 'Trang 1253-1257, Journal of Early Adolescence, Tập 44(9), 2023')

    H2(d, "2.1. Cohort z-proso")
    acro(d, "Dự án Phát triển Xã hội Zurich", "Zurich Project on Social Development from Childhood to Adulthood", "z-proso")
    nr(d, "• Mẫu ban đầu: 1.675 trẻ em vào lớp 1 tại 56 trường tiểu học công lập "
       "Zurich, năm 2004. Phương pháp lấy mẫu phân tầng ngẫu nhiên có TĂNG TỈ LỆ "
       "MẪU ở các quận khó khăn (oversampling of disadvantaged districts).")
    nr(d, "• Lịch trình đánh giá:")
    nr(d, "  - Tuổi 7 (n=1.360) đến tuổi 20 (2018, n=1.180): 8 đợt đánh giá")
    nr(d, "  - 4 đợt đánh giá thêm tháng 4-9/2020 (giai đoạn đại dịch, ~22 tuổi)")
    nr(d, "• Mẫu phân tích chính: n=786 (số người tham gia đợt khảo sát tháng 4/2020)")
    nr(d, "• Đợt đánh giá CHÍNH cho phân tích này: tuổi 13 (2011, n=1.365) so với "
       "tuổi ~22 (4/2020) — khoảng cách 9 năm")

    H2(d, "2.2. Phương pháp khảo sát + đạo đức")
    nr(d, "• Tuổi 13: Bảng câu hỏi giấy, làm trong lớp học (~90 phút)")
    nr(d, "• Năm 2020: Bảng câu hỏi online (~15-20 phút), thời gian hoàn thành 7 ngày")
    nr(d, "• Thưởng: ~30 USD ở tuổi 13; cơ hội trúng giải ~100 USD ở năm 2020")
    nr(d, "• Đồng thuận: Văn bản (tuổi 13, có quyền từ chối của cha mẹ); online (tuổi 22)")
    nr(d, "• Phê duyệt đạo đức: Hội đồng Đạo đức, Khoa Xã hội & Nhân văn, Đại học Zurich")
    nr(d, "• Có áp dụng trọng số mẫu (survey weighting) để suy luận cho mẫu gốc "
       "(Nivette và cộng sự 2021)")

    H2(d, "2.3. Thang đo (kết quả ở tuổi 22 — đại dịch)")

    H3(d, "2.3.1. Đo căng thẳng tâm lý (đầu đại dịch, 4/2020)")
    nr(d, "3 mục đo trên thang 10 điểm:")
    nr(d, "(1) \"Bạn cảm thấy TỆ HƠN hay TỐT HƠN từ khi đại dịch corona bắt đầu?\" "
       "(1=tệ hơn nhiều → 10=tốt hơn nhiều; mã hoá ngược: cao hơn = tệ hơn)")
    nr(d, "(2) \"Bạn cảm thấy LẠC QUAN về tương lai đến mức nào?\" (1=hoàn toàn không "
       "lạc quan → 10=rất lạc quan; mã hoá ngược: cao hơn = tuyệt vọng)")
    nr(d, "(3) \"Đại dịch đã LÀM XÁO TRỘN cuộc sống của bạn như thế nào, xét về thói "
       "quen, công việc, học tập, gia đình?\" (1=xáo trộn cao → 10=không xáo trộn; "
       "mã hoá ngược: cao hơn = xáo trộn nhiều)")

    H3(d, "2.3.2. Đo ứng phó thích nghi (đợt 4/2020)")
    nr(d, "Người tham gia báo cáo TẦN SUẤT sử dụng các chiến lược ứng phó trong 2 "
       "tuần qua khi gặp căng thẳng (thang 4 điểm: 1=không bao giờ → 4=rất thường "
       "xuyên). Các mục bao gồm:")
    nr(d, "• Mục đặc thù đại dịch: chấp nhận, đánh giá lại tích cực")
    nr(d, "• Mục điều chỉnh từ Carver 1997 (Brief COPE): tự phân tâm, tìm hỗ trợ cảm "
       "xúc, duy trì liên lạc")
    nr(d, "• Mục tự xây dựng: thói quen hàng ngày, tập thể dục, giúp đỡ hàng xóm, "
       "tìm hỗ trợ chuyên môn về SKTT")
    nr(d, "Cronbach α tổng = 0,54 (THẤP — phản ánh các chiến lược được đo trong các "
       "bối cảnh khác nhau).")
    nr(d, "Phương pháp tính tổng điểm: tổng tần suất sử dụng tất cả ứng phó thích "
       "nghi (phổ ~4-16, phản ánh sự đa dạng + tần suất trung bình).")
    nr(d, "Các thang con: (a) Ứng phó xã hội — tìm hỗ trợ cảm xúc, duy trì liên lạc, "
       "giúp đỡ người khác, tìm hỗ trợ chuyên môn; (b) Ứng phó tự túc — tự phân tâm, "
       "chấp nhận, đánh giá lại tích cực, thói quen hàng ngày, tập thể dục.")

    H2(d, "2.4. Thang đo (yếu tố tiên lượng ở tuổi 13)")

    H3(d, "2.4.1. Triệu chứng nội tâm hoá (Internalizing Symptoms)")
    nr(d, "8 mục từ Bảng câu hỏi Hành vi Xã hội (Social Behavior Questionnaire — SBQ) "
       "của Tremblay và cộng sự 1991: triệu chứng trầm cảm + lo âu (vd: \"Tôi cảm "
       "thấy buồn không lý do\"; \"Tôi cảm thấy lo âu\"). Thang 5 điểm (1=không bao "
       "giờ → 5=rất thường xuyên). Trung bình thang, α = 0,82 (cao hơn = nhiều "
       "triệu chứng hơn).")

    H3(d, "2.4.2. Sự kiện căng thẳng tích luỹ")
    nr(d, "Danh sách 25 sự kiện đời; 21 sự kiện được chọn là liên quan đến căng "
       "thẳng (chết người thân/thú cưng, cha mẹ ly thân, nhập viện). Tổng điểm: "
       "số sự kiện trải qua giữa tuổi 11-13.")

    H3(d, "2.4.3. Tương tác hỗ trợ cha-mẹ-con")
    acro(d, "Bảng câu hỏi nuôi dạy con (Alabama)", "Alabama Parenting Questionnaire", "APQ")
    nr(d, "6 mục từ APQ (Shelton và cộng sự 1996): hỗ trợ cảm xúc / thực tế, hoạt "
       "động chung (an ủi, giúp tìm trợ giúp, làm bài tập, nói chuyện / chơi cùng "
       "nhau). Thang 4 điểm (1=không bao giờ → 4=rất thường xuyên). Trung bình "
       "thang, α = 0,72 (cao hơn = hỗ trợ nhiều hơn).")

    H2(d, "2.5. Biến kiểm soát + bối cảnh")
    nr(d, "• SES: International Socio-Economic Index (ISEI), phổ 16-90, lấy điểm tối "
       "đa giữa tuổi 11-15, đảo ngược (cao hơn = SES thấp hơn)")
    nr(d, "• Nguồn gốc di cư cha mẹ: Nhị phân (1 = cả 2 cha mẹ sinh ở nước ngoài; "
       "0 = ≥1 sinh ở Thuỵ Sĩ)")
    nr(d, "• Giới tính: Nhị phân (1 = nữ; 0 = nam)")
    nr(d, "• Tác nhân căng thẳng tích luỹ liên quan đại dịch (4/2020): tổng 15 sự "
       "kiện căng thẳng (sức khỏe: phơi nhiễm COVID-19; kinh tế / giáo dục: mất việc, "
       "khó khăn tài chính)")

    H2(d, "2.6. Chiến lược phân tích")
    nr(d, "Phân tích chính: hồi quy tuyến tính bằng phần mềm MPlus phiên bản 8 với "
       "phương pháp ước lượng MLR (Maximum Likelihood with Robust standard errors) — "
       "robust với phân phối không chuẩn.")
    nr(d, "BƯỚC 1: Mô hình riêng cho từng cặp predictor-outcome (kiểm soát các biến "
       "nhân khẩu học)")
    nr(d, "BƯỚC 2: Mô hình đầy đủ (cả 3 predictor cùng lúc) + kiểm soát biến nhân "
       "khẩu học + tác nhân căng thẳng đại dịch")
    nr(d, "Outcomes: 3 chỉ báo căng thẳng tâm lý + tần suất ứng phó thích nghi")
    nr(d, "Tác động tương tác (interaction effects): Tương tác 2 chiều giữa từng "
       "predictor và tác nhân căng thẳng đại dịch lên căng thẳng tâm lý (đã căn "
       "giữa biến — centered)")
    nr(d, "Tác động gián tiếp (indirect effects): Mô hình đường dẫn (path models) — "
       "predictor ở tuổi 13 → ứng phó thích nghi → căng thẳng tâm lý ở đợt thứ 2 "
       "(cuối 4 / đầu 5/2020). Bootstrapped standard errors (1.000 lần lấy mẫu, "
       "bias-corrected). Hiệu ứng ứng phó lên căng thẳng tâm lý điều chỉnh cho "
       "biến nhân khẩu học + tác nhân đại dịch + căng thẳng tâm lý baseline.")
    nr(d, "Xử lý dữ liệu thiếu: 1-2% mất ở biến 4/2020; ~5% ở biến tuổi 13; 18% "
       "ở đợt cuối 4/đầu 5/2020 (chủ yếu do bỏ giữa). Kiểm định Little MCAR p = "
       "0,059 (không có ý nghĩa). Phương pháp: Multiple Imputation (Bayesian "
       "estimation, 20 datasets).")

    # ========== PHẦN 3 — KẾT QUẢ ==========
    H1(d, "PHẦN 3 — KẾT QUẢ (RESULTS)")
    page_marker(d, 'Trang 1257-1262, Journal of Early Adolescence, Tập 44(9), 2023')

    H2(d, "3.1. Thống kê mô tả mẫu (Bảng 1 — tái tạo, có trọng số)")
    tbl(d, ['Biến', 'Trung bình (M) hoặc %', 'Sai số chuẩn (SD)'], [
        ['SES thấp (tuổi 13, đảo ngược)', 'M = 43,05', 'SD = 20,02'],
        ['Cha mẹ di cư (cả 2)', '48% có', '—'],
        ['Nữ', '48%', '—'],
        ['Triệu chứng nội tâm hoá (tuổi 13)', 'M = 2,20', 'SD = 0,72'],
        ['Sự kiện căng thẳng tích luỹ (tuổi 11-13)', 'M = 2,66', 'SD = 1,82'],
        ['Tương tác hỗ trợ cha-mẹ-con (tuổi 13)', 'M = 3,09', 'SD = 0,56'],
        ['Cảm thấy tệ hơn (4/2020, tuổi ~22)', 'M = 5,80', 'SD = 1,48'],
        ['Tuyệt vọng', 'M = 4,12', 'SD = 2,03'],
        ['Cảm nhận xáo trộn cuộc sống', 'M = 6,28', 'SD = 2,44'],
        ['Tần suất ứng phó thích nghi (tổng)', 'M = 21,61', 'SD = 3,61'],
        ['Tần suất ứng phó thích nghi (trung bình)', 'M = 2,40', 'SD = 0,40'],
        ['Tác nhân căng thẳng đại dịch tích luỹ', 'M = 0,92', 'SD = 1,13'],
    ], [7.0, 4.0, 4.0])

    H2(d, "3.2. Tương quan chính (mức ý nghĩa thống kê)")
    nr(d, "• Triệu chứng nội tâm hoá ↔ tuyệt vọng: r = 0,14 (p<0,001)")
    nr(d, "• Triệu chứng nội tâm hoá ↔ xáo trộn cuộc sống: r = 0,12 (p<0,01)")
    nr(d, "• Tương tác hỗ trợ ↔ ứng phó thích nghi: r = 0,20 (p<0,001) — TƯƠNG QUAN MẠNH NHẤT")
    nr(d, "• Ứng phó thích nghi ↔ cảm thấy tệ hơn: r = −0,15 (p<0,001)")
    nr(d, "• Ứng phó thích nghi ↔ tuyệt vọng: r = −0,14 (p<0,001)")
    nr(d, "• Tác nhân đại dịch ↔ xáo trộn cuộc sống: r = 0,14 (p<0,001)")

    H2(d, "3.3. Yếu tố tiên lượng tuổi 13 → căng thẳng tâm lý tuổi 22 (Bảng 2)")

    H3(d, "3.3.1. Outcome: Cảm thấy tệ hơn (4/2020)")
    tbl(d, ['Yếu tố tiên lượng', 'Mô hình riêng (β, p)', 'Mô hình đầy đủ (β, p)'], [
        ['Triệu chứng nội tâm hoá', 'β=0,10, p=0,014*', 'β=0,09, p=0,042*'],
        ['Sự kiện căng thẳng', 'β=0,06, p=0,155 ns', 'β=0,04, p=0,405 ns'],
        ['Tương tác hỗ trợ', 'β=−0,03, p=0,487 ns', 'β=0,00, p=0,963 ns'],
        ['Tác nhân đại dịch', '—', 'β=0,03, p=0,422 ns'],
    ], [5.5, 5.0, 5.0])

    H3(d, "3.3.2. Outcome: Tuyệt vọng (4/2020)")
    tbl(d, ['Yếu tố tiên lượng', 'Mô hình riêng', 'Mô hình đầy đủ'], [
        ['Triệu chứng nội tâm hoá', 'β=0,10, p=0,010**', 'β=0,11, p=0,011*'],
        ['Sự kiện căng thẳng', 'β=−0,02, p=0,611 ns', 'β=−0,05, p=0,256 ns'],
        ['Tương tác hỗ trợ', 'β=−0,06, p=0,188 ns', 'β=−0,04, p=0,387 ns'],
        ['Nữ giới', 'β=0,13, p<0,001***', 'β=0,10, p=0,007**'],
        ['Tác nhân đại dịch', '—', 'β=−0,01, p=0,882 ns'],
    ], [5.5, 5.0, 5.0])

    H3(d, "3.3.3. Outcome: Xáo trộn cuộc sống (4/2020)")
    tbl(d, ['Yếu tố tiên lượng', 'Mô hình riêng', 'Mô hình đầy đủ'], [
        ['Triệu chứng nội tâm hoá', 'β=0,09, p=0,020*', 'β=0,09, p=0,027*'],
        ['Sự kiện căng thẳng', 'β=0,03, p=0,507 ns', 'β=0,00, p=0,976 ns'],
        ['Tương tác hỗ trợ', 'β=0,04, p=0,320 ns', 'β=0,08, p=0,062 (gần ý nghĩa)'],
        ['Nữ giới', 'β=0,12, p=0,001**', 'β=0,09, p=0,023*'],
        ['SES thấp', 'β=−0,10, p=0,026*', 'β=−0,09, p=0,056 (gần ý nghĩa)'],
        ['Tác nhân đại dịch', 'β=0,15, p<0,001***', 'β=0,15, p<0,001***'],
    ], [5.5, 5.0, 5.0])

    nr(d, "PHÁT HIỆN 1 — Triệu chứng nội tâm hoá tuổi 13 → căng thẳng tâm lý tuổi 22:",
       bold=True, color=BLUE)
    nr(d, "Triệu chứng nội tâm hoá ở tuổi 13 NHẤT QUÁN tiên lượng cả 3 chỉ báo căng "
       "thẳng tâm lý 9 năm sau (β = 0,09-0,11; p < 0,05). Đây là bằng chứng vững "
       "chắc cho \"liên tục SKTT\" theo thời gian (Copeland và cộng sự 2009).",
       bold=True)

    H2(d, "3.4. Tác động tương tác — Bằng chứng cho giả thuyết MIỄN DỊCH TÂM LÝ")
    nr(d, "Tác động tương tác có ý nghĩa duy nhất: Sự kiện căng thẳng tích luỹ × Tác "
       "nhân đại dịch lên xáo trộn cuộc sống: β = −0,11, p = 0,003**.")
    nr(d, "DIỄN GIẢI: Tác động của tác nhân đại dịch lên cảm nhận xáo trộn GIẢM khi "
       "số sự kiện căng thẳng ở tuổi 13 TĂNG. Trong số những người trải qua ≥1,5 sự "
       "kiện căng thẳng trên trung bình ở tuổi vị thành niên sớm, tác nhân đại dịch "
       "KHÔNG còn liên quan có ý nghĩa với xáo trộn (KTC 95% bao gồm 0).")
    nr(d, "→ ỦNG HỘ GIẢ THUYẾT MIỄN DỊCH TÂM LÝ (Inoculation Hypothesis): trải qua "
       "khó khăn ở mức độ vừa phải khi còn vị thành niên có thể \"tiêm vaccine\" cho "
       "khả năng đối phó với khủng hoảng lớn sau này. Đây là phát hiện ĐẶC BIỆT QUAN "
       "TRỌNG và phản trực giác.", bold=True, color=BLUE)
    nr(d, "Các tác động tương tác khác: KHÔNG có ý nghĩa thống kê.")

    H2(d, "3.5. Yếu tố tiên lượng tuổi 13 → ứng phó thích nghi tuổi 22")
    tbl(d, ['Yếu tố tiên lượng', 'Mô hình riêng', 'Mô hình đầy đủ'], [
        ['Triệu chứng nội tâm hoá', 'β=−0,06, p=0,115 ns', 'β=−0,04, p=0,311 ns'],
        ['Sự kiện căng thẳng', 'β=0,01, p=0,783 ns', 'β=0,03, p=0,511 ns'],
        ['Tương tác hỗ trợ', 'β=0,14, p<0,001***', 'β=0,14, p=0,001**'],
        ['Nữ giới', 'β=0,20, p<0,001***', 'β=0,20, p<0,001***'],
        ['SES thấp', 'β=−0,18, p<0,001***', 'β=−0,15, p<0,001***'],
        ['Tác nhân đại dịch', '—', 'β=0,05, p=0,293 ns'],
    ], [5.5, 5.0, 5.0])

    nr(d, "PHÁT HIỆN 2 — TƯƠNG TÁC HỖ TRỢ tuổi 13 → ứng phó thích nghi tuổi 22:",
       bold=True, color=BLUE)
    nr(d, "Tương tác hỗ trợ cha-mẹ-con tuổi 13 NHẤT QUÁN tiên lượng tần suất ứng "
       "phó thích nghi 9 năm sau (β = 0,14, p < 0,001). Đặc biệt, hiệu ứng còn KÉO "
       "DÀI sau khi kiểm soát biến nhân khẩu học + tác nhân đại dịch — chứng tỏ "
       "tác động ĐỘC LẬP và LÂU DÀI.", bold=True)
    nr(d, "Phát hiện thú vị: Triệu chứng nội tâm hoá KHÔNG tiên lượng việc DÙNG ÍT "
       "ứng phó thích nghi — gợi ý vị thành niên có triệu chứng nội tâm hoá VẪN có "
       "thể dùng ứng phó thích nghi (đồng thời cũng dùng ứng phó không thích nghi).")

    H2(d, "3.6. Chiến lược ứng phó cụ thể nào được tương tác hỗ trợ tiên lượng?")
    nr(d, "Phân tích chi tiết từng chiến lược (Bảng 3):")
    nr(d, "ỨNG PHÓ XÃ HỘI:")
    nr(d, "• Tìm hỗ trợ cảm xúc: β = 0,11 (p=0,002**)")
    nr(d, "• Duy trì liên lạc với gia đình + bạn bè: β = 0,11 (p=0,012*)")
    nr(d, "• Giúp đỡ người khác: β = 0,11 (p=0,006**)")
    nr(d, "• Tìm hỗ trợ chuyên môn: tác nhân đại dịch β = 0,14 (p=0,024*); nữ — bảo "
       "vệ β = −0,08 (p=0,042*)")
    nr(d, "ỨNG PHÓ TỰ TÚC:")
    nr(d, "• Tự phân tâm: nữ β = 0,17 (p<0,001***); tác nhân đại dịch β = 0,11 (p=0,011*)")
    nr(d, "• Chấp nhận: SES thấp β = −0,13 (p=0,003**)")
    nr(d, "• Đánh giá lại nhận thức: tương tác hỗ trợ β = 0,10 (p=0,019*); SES thấp "
       "β = −0,10 (p=0,017*)")
    nr(d, "• Thói quen hàng ngày: nữ β = 0,13 (p=0,001**); tác nhân đại dịch — bảo "
       "vệ β = −0,10 (p=0,006**)")
    nr(d, "• Hoạt động thể chất: SES thấp β = −0,18 (p<0,001***); cha mẹ di cư β = "
       "−0,09 (p=0,028*)")
    nr(d, "TÓM TẮT THANG CON:")
    nr(d, "• Ứng phó tự túc: tương tác hỗ trợ β = 0,09 (p=0,030*)")
    nr(d, "• Ứng phó xã hội: tương tác hỗ trợ β = 0,15 (p<0,001***) — TÁC ĐỘNG MẠNH NHẤT")

    H2(d, "3.7. Tác động gián tiếp: tuổi 13 → ứng phó → căng thẳng tâm lý")
    nr(d, "Tác động trực tiếp của ứng phó thích nghi lên căng thẳng tâm lý sau (đợt "
       "cuối 4/đầu 5/2020, n=694):")
    nr(d, "• Cảm thấy tệ hơn: β = −0,10 (p=0,029*)")
    nr(d, "• Tuyệt vọng: β = −0,11 (p=0,011*)")
    nr(d, "• Xáo trộn cuộc sống: β = −0,03 (p=0,426 ns)")
    nr(d, "Tác động GIÁN TIẾP của tương tác hỗ trợ tuổi 13 lên căng thẳng tâm lý "
       "qua ứng phó:")
    nr(d, "• Cảm thấy tệ hơn: b không chuẩn hoá = −0,05 (95% CI: −0,11 đến −0,01)*")
    nr(d, "• Tuyệt vọng: b = −0,06 (95% CI: −0,12 đến −0,02)*")
    nr(d, "• Xáo trộn cuộc sống: không có ý nghĩa")

    nr(d, "PHÁT HIỆN 3 — Đường dẫn gián tiếp được xác lập:", bold=True, color=BLUE)
    nr(d, "Tương tác hỗ trợ cha-mẹ-con tuổi 13 → tăng ứng phó thích nghi tuổi 22 → "
       "GIẢM cảm giác tệ hơn + tuyệt vọng trong đại dịch. Đây là MÔ HÌNH 3 BƯỚC THEO "
       "THỜI GIAN, kiểm chứng được nhờ thiết kế dài hạn.", bold=True)

    # ========== PHẦN 4 — THẢO LUẬN ==========
    H1(d, "PHẦN 4 — THẢO LUẬN (DISCUSSION)")
    page_marker(d, 'Trang 1262-1268, Journal of Early Adolescence, Tập 44(9), 2023')

    H2(d, "4.1. Vai trò của triệu chứng nội tâm hoá tuổi vị thành niên sớm")
    nr(d, "Triệu chứng nội tâm hoá tuổi vị thành niên sớm NHẤT QUÁN tiên lượng MỨC ĐỘ "
       "CAO HƠN của cả 3 chỉ báo căng thẳng tâm lý ~9 năm sau trong COVID-19. Phản "
       "ánh: (1) sự liên tục của SKTT theo thời gian (Copeland 2009; Rao & Chen "
       "2009), và (2) độ dễ tổn thương cao hơn của những người có khuynh hướng "
       "nội tâm hoá khi gặp tác nhân căng thẳng mới.")
    nr(d, "PHÁT HIỆN CHƯA GIẢI THÍCH: Triệu chứng nội tâm hoá KHÔNG liên quan với "
       "GIẢM ứng phó thích nghi. Có 3 cách giải thích:")
    nr(d, "(1) Vị thành niên có vấn đề nội tâm hoá có thể dùng CẢ chiến lược không "
       "thích nghi (rumination, tự hại) VÀ thích nghi đồng thời — tạo hiệu ứng net "
       "= 0 (Krause và cộng sự 2018; Nock và cộng sự 2006)")
    nr(d, "(2) Một số có thể đã tiếp cận dịch vụ SKTT, phát triển điều hoà cảm xúc "
       "lành mạnh dù từng có triệu chứng (Wu và cộng sự 2001)")
    nr(d, "(3) Tính không đồng nhất nhóm con (người được hỗ trợ vs không) có thể "
       "tạo các liên hệ đối kháng nhau, triệt tiêu lẫn nhau")

    H2(d, "4.2. Vai trò của sự kiện căng thẳng tuổi vị thành niên sớm")
    nr(d, "Bằng chứng ỦNG HỘ GIẢ THUYẾT MIỄN DỊCH TÂM LÝ. Thanh niên có phơi nhiễm "
       "căng thẳng cao bất thường ở tuổi vị thành niên sớm KHÔNG có nguy cơ tăng "
       "xáo trộn cuộc sống khi đối mặt với tác nhân đại dịch. Cơ chế giả định: cách "
       "khung nhận thức khác nhau — người trải qua nhiều sự kiện căng thẳng trước "
       "đó có thể có ngưỡng cao hơn để xem các sự kiện thêm là \"xáo trộn\" cuộc "
       "sống.")
    nr(d, "Lưu ý: KHÔNG có tác động chính trên tuyệt vọng hay cảm thấy tệ hơn, và "
       "KHÔNG có tương tác với tác nhân đại dịch cho 2 outcome đó.")
    nr(d, "Hạn chế: Đo lường tập trung vào MERE EXPOSURE đến sự kiện căng thẳng — "
       "không đánh giá cách người tham gia ĐÁNH GIÁ và VƯỢT QUA chúng. Người vượt "
       "qua thành công được hưởng lợi từ inoculation; người không thành công vẫn "
       "dễ tổn thương. Nghiên cứu tương lai cần thông tin bổ sung về CHẤT LƯỢNG "
       "vượt qua (Oldehinkel 2014).")

    H2(d, "4.3. Vai trò của tương tác hỗ trợ cha-mẹ-con")
    nr(d, "Tương tác hỗ trợ cha-mẹ-con KHÔNG liên quan TRỰC TIẾP với căng thẳng tâm "
       "lý ở tuổi 22 — nhưng GIÁN TIẾP hỗ trợ phát triển ứng phó thích nghi → giảm "
       "căng thẳng tâm lý.")
    nr(d, "3 cơ chế ứng phó cụ thể:")
    nr(d, "(1) CHIẾN LƯỢC XÃ HỘI: Liên hệ TRỰC TIẾP MẠNH (β=0,11-0,15) với tìm hỗ "
       "trợ, duy trì liên lạc, giúp đỡ người khác. Cha mẹ làm gương + dạy rằng hỗ "
       "trợ xã hội là sẵn có và có giá trị (Skinner & Zimmer-Gembeck 2007). Kỹ năng "
       "này chuyển sang quan hệ bạn bè / người yêu ở tuổi trưởng thành.")
    nr(d, "(2) CHIẾN LƯỢC NHẬN THỨC: Liên hệ với đánh giá lại nhận thức (β=0,10) gợi "
       "ý cha mẹ rèn luyện khả năng \"đặt mình vào vị trí người khác\" / nhìn nhận "
       "đa chiều (Hall và cộng sự 2021), giúp con khung lại các tác nhân căng thẳng "
       "theo hướng tích cực.")
    nr(d, "(3) CHIẾN LƯỢC HOẠT ĐỘNG: Có thể liên quan đến chiến lược tự túc (β=0,09, "
       "gần ý nghĩa) qua việc cha mẹ làm gương các hoạt động chung — tăng mức độ "
       "hoạt động chung và tham gia hoạt động giải trí.")
    nr(d, "ĐÓNG GÓP MỚI: Hầu hết nghiên cứu trước xem xét hiệu quả ứng phó SAU KHI "
       "đã có; bài này cho thấy XÃ HỘI HOÁ GIA ĐÌNH ở tuổi vị thành niên sớm hình "
       "thành CAPACITY ứng phó SUỐT ĐỜI. Trong giai đoạn phong toả, các thành viên "
       "trong gia đình trở thành nguồn hỗ trợ xã hội CHÍNH; do đó kỹ năng ứng phó "
       "xã hội được học sớm (từ xã hội hoá gia đình) tỏ ra ĐẶC BIỆT HỮU ÍCH.",
       bold=True, color=BLUE)

    H2(d, "4.4. Phát hiện bổ sung")
    nr(d, "• KHÁC BIỆT GIỚI: Nữ báo cáo nhiều tuyệt vọng, xáo trộn cuộc sống, và "
       "dùng ứng phó thích nghi NHIỀU HƠN nam.")
    nr(d, "• HIỆU ỨNG SES: SES thấp ở tuổi vị thành niên liên quan với tần suất "
       "ứng phó thích nghi THẤP HƠN + giảm engagement vào các chiến lược cụ thể "
       "(chấp nhận, đánh giá lại nhận thức, hoạt động thể chất).")
    nr(d, "• KÍCH THƯỚC HIỆU ỨNG nhỏ NHƯNG đáng chú ý do khoảng cách 9 năm. Götz "
       "và cộng sự (2022) lập luận các kích thước hiệu ứng nhỏ này là nền tảng "
       "cho khoa học tâm lý tích luỹ.")

    # ========== PHẦN 5 — DANH MỤC TÀI LIỆU THAM KHẢO ==========
    H1(d, "PHẦN 5 — DANH MỤC TÀI LIỆU THAM KHẢO TRONG BÀI GỐC")
    nr(d, "Bài gốc Steinhoff và cộng sự (2023) có TỔNG CỘNG 84 tài liệu tham khảo. "
       "Em liệt kê 20 tài liệu đầu — phần em đã đối chiếu trực tiếp từ PMC NCBI.",
       italic=True, color=GRAY)
    nr(d, "(Giữ nguyên tiếng Anh — định dạng gốc của tạp chí Journal of Early "
       "Adolescence theo APA)", italic=True, color=GRAY)

    refs2 = [
        "Aldwin, C. M. (2012). Stress and coping across the lifespan. In S. Folkman (Ed.), The Oxford Handbook of Stress, Health, and Coping. Oxford University Press.",
        "Appleyard, K., Egeland, B., Van Dulmen, M. H. M., & Sroufe, L. A. (2005). When more is not better: The role of cumulative risk in child behavior. Journal of Child Psychology and Psychiatry, 46(3), 235–245. https://doi.org/10.1111/j.1469-7610.2004.00351.x",
        "Berg, N., Kiviruusu, O., Karvonen, S., Rahkonen, O., & Huurre, T. (2017). Pathways from problems in adolescent family relationships to midlife mental health via early adulthood disadvantages. PLoS ONE, 12(5), e0178136.",
        "Buckley, L., Broadley, M., & Cascio, C. N. (2019). Socio-economic status and the developing brain in adolescence: A systematic review. Child Neuropsychology, 25(7), 859–884.",
        "Buhle, J. T., Silvers, J. A., Wager, T. D., et al. (2013). Cognitive reappraisal of emotion: A meta-analysis of human neuroimaging studies. Cerebral Cortex, 24(11), 2981–2990.",
        "Calvete, E., Orue, I., & Hankin, B. L. (2015). Cross-lagged associations among ruminative response style, stressors, and depressive symptoms in adolescents. Journal of Social and Clinical Psychology, 34(3), 203–220.",
        "Carroll, L. (2013). Problem-focused coping. In M. D. Gellman & J. R. Turner (Eds.), Encyclopedia of Behavioral Medicine (pp. 1540–1541). Springer.",
        "Carver, C. S. (1997). You want to measure coping but your protocol's too long: Consider the brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100.",
        "Chen, X., Zou, Y., & Gao, H. (2021). Role of neighborhood social support in stress coping and psychological wellbeing during the COVID-19 pandemic. Health & Place, 69, 102532.",
        "Cheng, C., Lau, H.-P. B., & Chan, M.-P. S. (2014). Coping flexibility and psychological adjustment to stressful life changes: A meta-analytic review. Psychological Bulletin, 140(6), 1582–1607.",
        "Compas, B. E., Jaser, S. S., Bettis, A. H., et al. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991.",
        "Conner, K. R., Duberstein, P. R., Conwell, Y., Seidlitz, L., & Caine, E. D. (2001). Psychological vulnerability to completed suicide: A review of empirical studies. Suicide and Life-Threatening Behavior, 31(4), 367–385.",
        "Copeland, W. E., Shanahan, L., Costello, E. J., & Angold, A. (2009). Childhood and adolescent psychiatric disorders as predictors of young adult disorders. Archives of General Psychiatry, 66(7), 764–772.",
        "Enders, C. K. (2013). Dealing with missing data in developmental research. Child Development Perspectives, 7(1), 27–31.",
        "Eschenbeck, H., Schmid, S., Schröder, I., et al. (2018). Development of coping strategies from childhood to adolescence. European Journal of Health Psychology, 25(1), 18–30.",
        "Fegert, J. M., Vitiello, B., Plener, P. L., & Clemens, V. (2020). Challenges and burden of the coronavirus 2019 (COVID-19) pandemic for child and adolescent mental health. Child and Adolescent Psychiatry and Mental Health, 14, 20.",
        "Folkman, S. (2010). Stress, coping, and hope. Psycho-Oncology, 19(9), 901–908.",
        "Folkman, S. (Ed.). (2012). The Oxford Handbook of Stress, Health, and Coping. Oxford University Press.",
        "Ford, C. A., Pool, A. C., Kahn, N. F., Jaccard, J., & Halpern, C. T. (2023). Associations between mother-adolescent and father-adolescent relationships and young adult health. JAMA Network Open, 6(3), e233944.",
        "Ganzeboom, H. B. G., De Graaf, P. M., Treiman, D. J., & De Leeuw, J. (1992). A standard international socio-economic index of occupational status. Social Science Research, 21, 1–56.",
    ]
    for i, r in enumerate(refs2, 1):
        nr(d, f"{i}. {r}", size=11)
    nr(d, "(Còn 64 tài liệu tham khảo tiếp theo — thầy có thể truy cập đầy đủ tại "
       "PMC10261967)", italic=True, color=GRAY, size=10)

    # ========== PHẦN 6 — PHẢN BIỆN ==========
    H1(d, "PHẦN 6 — PHẢN BIỆN CHI TIẾT (CRITICAL REVIEW)")

    H2(d, "Điểm 1 — Thiết kế cohort dài hạn 9 năm là điểm mạnh QUAN TRỌNG NHẤT")
    crit_para(d, "Theo dõi cùng một mẫu 786 người trong 9 năm — từ tuổi 13 (2011) đến "
              "tuổi 22 (2020 đại dịch) — là thiết kế HIẾM CÓ trong nghiên cứu coping. "
              "Cho phép suy luận thời gian (temporal inference): yếu tố tuổi 13 ĐI "
              "TRƯỚC outcome tuổi 22 → loại trừ được phần nào hướng nhân quả ngược "
              "(reverse causation) đã làm khó hiểu phát hiện của Herres 2015 cắt "
              "ngang. Tỉ lệ giữ mẫu (retention) tốt qua 12 đợt đánh giá — bằng chứng "
              "uy tín của z-proso.")

    H2(d, "Điểm 2 — Bằng chứng cho INOCULATION HYPOTHESIS đáng chú ý")
    crit_para(d, "Phát hiện rằng thanh niên có tích luỹ sự kiện căng thẳng cao ở "
              "tuổi 13 KHÔNG có nguy cơ tăng xáo trộn cuộc sống khi đối mặt với "
              "đại dịch (β tương tác = −0,11, p < 0,01) là CHỨNG CỨ ĐÁNG CHÚ Ý cho "
              "giả thuyết miễn dịch tâm lý. Lưu ý: hiệu ứng chỉ ở 1 trong 3 outcome "
              "(xáo trộn cuộc sống), KHÔNG ở tuyệt vọng hay cảm thấy tệ hơn — gợi ý "
              "miễn dịch chỉ ở khía cạnh nhận thức về xáo trộn, không phải emotion. "
              "Cần thận trọng khái quát hoá — không phải mọi căng thẳng sớm đều có "
              "lợi.")

    H2(d, "Điểm 3 — Effect sizes nhỏ — ý nghĩa thực tiễn cần xem xét")
    crit_para(d, "Hầu hết β = 0,09-0,15 — nhỏ theo chuẩn Cohen (β nhỏ < 0,10, "
              "trung bình ~0,30). Tác giả tự thừa nhận. Tuy nhiên Götz và cộng sự "
              "(2022) lập luận hiệu ứng nhỏ qua khoảng thời gian dài (9 năm) là "
              "đáng chú ý cho khoa học tâm lý tích luỹ. Áp dụng thực tiễn: KHÔNG "
              "kỳ vọng can thiệp \"hỗ trợ cha mẹ\" sẽ tạo thay đổi LỚN ở từng cá "
              "nhân — nhưng tích luỹ qua quần thể có thể có giá trị y tế công cộng.")

    H2(d, "Điểm 4 — Thang đo ứng phó hạn chế — α = 0,54")
    crit_para(d, "Tác giả tự thừa nhận \"limited\" adaptive coping measure: chỉ vài "
              "mục, hệ số tin cậy nội tại tổng α = 0,54 (THẤP). Do thang đo bao quát "
              "nhiều chiến lược trong các bối cảnh khác nhau (cá nhân, xã hội, "
              "đặc thù đại dịch), heterogeneity cao là dự kiến. Tuy nhiên độ tin "
              "cậy thấp làm GIẢM khả năng phát hiện hiệu ứng (statistical power) — "
              "hiệu ứng thực có thể MẠNH HƠN số liệu báo cáo. Nghiên cứu tương lai "
              "nên dùng thang đo dài hơn (vd. Brief COPE 28 mục đầy đủ).")

    H2(d, "Điểm 5 — Generalisability cho Việt Nam hạn chế")
    crit_para(d, "Mẫu Thuỵ Sĩ với SES trung bình - cao, hệ thống y tế + phúc lợi tốt; "
              "30% chuyển khỏi nhà cha mẹ ở tuổi 22 (đa số sống cùng bạn / bạn đời). "
              "Việt Nam có cấu trúc gia đình rất khác — đa thế hệ, sống cùng cha mẹ "
              "đến lập gia đình, áp lực xã hội về \"hiếu thảo\". Tác động của tương "
              "tác hỗ trợ cha-mẹ-con có thể MẠNH HƠN ở Việt Nam (do gia đình quan "
              "trọng hơn) hoặc YẾU HƠN (do không phải lựa chọn tự nguyện). Cần "
              "nghiên cứu nhân rộng (replication) ở Việt Nam.")

    # ========== PHẦN 7 — GAP ==========
    H1(d, "PHẦN 7 — PHÁT HIỆN GAP NGHIÊN CỨU TỪ BÀI NÀY")

    H2(d, "Gap 1 — Việt Nam chưa có nghiên cứu cohort dài hạn về SKTT thanh thiếu niên")
    nr(d, "Z-proso (Thuỵ Sĩ) theo dõi 9 năm. ALSPAC (Anh) theo dõi 30 năm. Đại học "
       "Queensland (Úc) Cohort theo dõi 14 năm. Việt Nam: chưa có cohort dài hạn "
       "tương đương. Đây là khoảng trống nghiên cứu LỚN — cần Bộ KHCN / Quỹ NAFOSTED "
       "đầu tư.")

    H2(d, "Gap 2 — Thiếu thang đo Alabama Parenting Questionnaire (APQ) bản Việt")
    nr(d, "APQ là thang đo chuẩn quốc tế đo nuôi dạy con (parenting practices). Chưa "
       "có phiên bản Việt được kiểm định. Cần dịch + back-translate + kiểm định cho "
       "mẫu Việt.")

    H2(d, "Gap 3 — Việt Nam chưa kiểm chứng giả thuyết MIỄN DỊCH TÂM LÝ")
    nr(d, "Phát hiện inoculation ở Thuỵ Sĩ có lặp lại ở Việt Nam không? Đặc biệt "
       "trong context Việt Nam có nhiều khó khăn lịch sử (chiến tranh, hậu chiến, "
       "đói nghèo, COVID-19, kinh tế suy thoái). Có thể là chủ đề nghiên cứu rất "
       "có giá trị.")

    H2(d, "Gap 4 — Vai trò ÔNG BÀ trong tương tác hỗ trợ — chưa được Steinhoff đo")
    nr(d, "Bài Steinhoff chỉ đo cha mẹ. Văn hoá Việt Nam có vai trò ÔNG BÀ trong "
       "nuôi dạy cháu rất mạnh — đặc biệt khi cha mẹ đi làm xa. Nghiên cứu Việt Nam "
       "cần đo CẢ tương tác cháu-ông-bà song song với cha-mẹ-con.")

    H2(d, "Gap 5 — Thiếu nghiên cứu can thiệp DỰA TRÊN bằng chứng này")
    nr(d, "Nếu tương tác hỗ trợ tuổi 13 → ứng phó thích nghi tuổi 22, vậy can thiệp "
       "ĐÀO TẠO CHA MẸ ở tuổi vị thành niên sớm (lớp 6-8) có cải thiện ứng phó tuổi "
       "thanh niên không? Cần thử nghiệm có đối chứng ngẫu nhiên (RCT) — chưa ai "
       "làm ở Việt Nam.")

    H2(d, "Gap 6 — Tương tác giữa SES + nuôi dạy + outcomes")
    nr(d, "Steinhoff phát hiện SES thấp → ứng phó thích nghi thấp. Nhưng SES + nuôi "
       "dạy có TƯƠNG TÁC nào không? Có thể nuôi dạy hỗ trợ BUFFER tác động xấu của "
       "SES thấp. Cần phân tích tương tác chi tiết hơn — đặc biệt quan trọng ở "
       "Việt Nam có khác biệt SES rõ rệt giữa thành thị-nông thôn-DTTS.")

    out = os.path.join(OUT_DIR, 'Steinhoff_2023_LongitudinalCoping_v3_chi_tiet_27042026.docx')
    d.save(out); return out

if __name__ == '__main__':
    p1 = make_doc_herres_v3(); print('Herres V3:', p1)
    p2 = make_doc_steinhoff_v3(); print('Steinhoff V3:', p2)
