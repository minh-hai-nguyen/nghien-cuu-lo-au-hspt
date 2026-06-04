# -*- coding: utf-8 -*-
"""Tạo file Word trả lời câu hỏi dịch thuật ngữ cho thầy Minh Đức."""
import os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/Dịch thuật ngữ tiếng Anh sang tiếng Việt.docx"
BLUE = RGBColor(0x00, 0x00, 0xCC)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(13)


def heading(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(13)
    return p


def body(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def answer(label, t):
    """Dòng đề xuất dịch — tô xanh."""
    p = doc.add_paragraph()
    r0 = p.add_run(label + " ")
    r0.bold = True
    r = p.add_run(t)
    r.bold = True
    r.font.color.rgb = BLUE
    return p


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("TRẢ LỜI: DỊCH CÁC THUẬT NGỮ TIẾNG ANH SANG TIẾNG VIỆT")
tr.bold = True
tr.font.size = Pt(14)

body("Kính gửi thầy Minh Đức,")
body("Dưới đây là đề xuất dịch và cách xử lý sáu thuật ngữ thầy nêu, để đưa vào "
     "bản thảo bài báo tiếng Việt. Phần in xanh là cụm thầy có thể dùng trực tiếp.")
doc.add_paragraph()

# 1
heading("1. Early Adolescents' Coping Power")
body("Thầy đề xuất dịch là “Năng lực ứng phó của tuổi đầu vị thành niên”. Về mặt "
     "nghĩa, cách dịch này chấp nhận được. Tuy nhiên đây là TÊN một chương trình "
     "can thiệp tâm lý (cùng loại với “Coping Cat”, “Journey of the Brave”), nên "
     "thông lệ là giữ tên gốc tiếng Anh và kèm phần dịch trong ngoặc để người đọc "
     "tra cứu được tài liệu gốc.")
body("Một lưu ý nhỏ về chữ “của”: chương trình này hướng ĐẾN nhóm tuổi đầu vị "
     "thành niên, nên dùng “cho/dành cho” sẽ chính xác hơn “của”.")
answer("→ Đề xuất dùng:",
       "chương trình Coping Power dành cho tuổi đầu vị thành niên "
       "(Early Adolescent Coping Power)")
body("Nếu thầy muốn dịch hẳn không giữ tiếng Anh, có thể dùng: “Năng lực Ứng phó "
     "cho Tuổi đầu Vị thành niên”.")
doc.add_paragraph()

# 2
heading("2. canonical_index.json")
body("Đây là tên một TỆP DỮ LIỆU nội bộ của nhóm — không nên để tên tệp tin trong "
     "bài báo khoa học.")
answer("→ Thay bằng:", "cơ sở dữ liệu tài liệu nội bộ")
body("(hoặc “tệp chỉ mục tài liệu nội bộ do nhóm tác giả xây dựng”).")
doc.add_paragraph()

# 3
heading("3. PubMed, ScienceDirect, Scopus")
body("Đây là tên riêng của ba cơ sở dữ liệu học thuật quốc tế. Theo ý thầy, gộp "
     "lại thành một cụm chung bằng tiếng Việt.")
answer("→ Thay bằng:", "các nền tảng cơ sở dữ liệu điện tử")
doc.add_paragraph()

# 4-5
heading("4. Child Psychiatry and Human Development  &  5. Nature Human Behaviour")
body("Đây là TÊN TẠP CHÍ khoa học. Theo chuẩn trích dẫn quốc tế (APA phiên bản "
     "thứ bảy), tên tạp chí GIỮ NGUYÊN tiếng Anh trong Danh mục tài liệu tham "
     "khảo, KHÔNG dịch sang tiếng Việt (nếu dịch sẽ không tra cứu được).")
answer("→ Trong Danh mục tài liệu tham khảo:",
       "giữ nguyên “Child Psychiatry & Human Development”, “Nature Human Behaviour”.")
answer("→ Trong thân bài:",
       "không cần nêu tên tạp chí — chỉ cần ghi tên tác giả và năm; tên tạp chí đã "
       "có đầy đủ trong Danh mục tài liệu tham khảo.")
doc.add_paragraph()

# 6
heading("6. MSPSS")
body("MSPSS là viết tắt của Multidimensional Scale of Perceived Social Support — "
     "tên một thang đo. Theo quy định của Tạp chí (không viết tắt), cần viết đầy "
     "đủ tên tiếng Việt ở lần dùng đầu tiên.")
answer("→ Lần dùng đầu tiên:",
       "Thang đo Đa chiều về Hỗ trợ Xã hội được Cảm nhận "
       "(Multidimensional Scale of Perceived Social Support – MSPSS)")
body("Các lần nhắc sau có thể dùng gọn: “thang đo MSPSS” hoặc “thang đo này”.")
doc.add_paragraph()

body("Trân trọng.")

cp = doc.core_properties
cp.author = ""
cp.last_modified_by = ""
doc.save(OUT)
print("Đã lưu:", OUT)
