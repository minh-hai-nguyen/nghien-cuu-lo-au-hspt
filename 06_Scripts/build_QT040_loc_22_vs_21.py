"""Build doc: QT040 Walder 2025 - quy trinh loc 22 vs 21."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def add_heading(text, level=1, color=(31, 73, 125)):
    h = d.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(*color)
    return h

def add_para(text, bold=False, color=None, size=11):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.font.size = Pt(size)
    r.bold = bold
    if color:
        r.font.color.rgb = RGBColor(*color)
    return p

title = d.add_heading('Quy trình lọc nghiên cứu của Walder et al. (2025)', level=0)
for run in title.runs:
    run.font.color.rgb = RGBColor(31, 73, 125)

sub = d.add_paragraph()
sub_r = sub.add_run('Meta-analysis về DMHI cho lo âu xã hội ở trẻ em, vị thành niên, thanh niên')
sub_r.italic = True
sub_r.font.size = Pt(12)
sub_r.font.color.rgb = RGBColor(90, 90, 90)

d.add_paragraph()

meta_tbl = d.add_table(rows=1, cols=1)
meta_tbl.style = 'Table Grid'
cell = meta_tbl.rows[0].cells[0]
shade_cell(cell, 'FFF8DC')
meta_p = cell.paragraphs[0]
meta_p.add_run('Nguồn: ').bold = True
meta_p.add_run('QT040 — Walder N, Frey A, Berger T et al. (2025). "Digital mental health interventions for prevention and treatment of social anxiety disorder in children, adolescents and young adults: Systematic review and meta-analysis." ')
r_j = meta_p.add_run('Journal of Medical Internet Research (JMIR, Q1, IF ≈ 7,4). ')
r_j.italic = True
meta_p.add_run('DOI: 10.2196/preprints.67067 | PROSPERO CRD42023424181.')

d.add_paragraph()

add_heading('1. Từ "hàng trăm" → 22 bài (Systematic Review)', level=1)
add_para('Thực tế khởi đầu là 2.149 records chứ không phải vài trăm. Chú ý: các con số 1.357 và 205 là số bài BỊ LOẠI tại mỗi bước sàng lọc (không phải số còn lại).')

d.add_paragraph()

tbl = d.add_table(rows=8, cols=3)
tbl.style = 'Table Grid'

headers = ['Bước', 'Thao tác', 'Số bài']
hdr = tbl.rows[0]
for i, h in enumerate(headers):
    cell = hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

flow_rows = [
    ('1', 'Tìm kiếm databases điện tử', '2.149 records'),
    ('2', 'Loại duplicates (trùng nhau)', '−474 → còn 1.675'),
    ('3', 'Sàng lọc TITLE (2 reviewer độc lập)', 'loại 1.357 → còn 318'),
    ('4', 'Sàng lọc ABSTRACT', 'loại 205 → còn 113'),
    ('5', 'Đánh giá FULL TEXT', '113 bài được đánh giá'),
    ('6', 'Đạt tiêu chí eligibility → vào SR', '22 bài'),
    ('7', 'Đủ dữ liệu post-test → vào MA', '21 bài'),
]
for i, (b, t, n) in enumerate(flow_rows, start=1):
    row = tbl.rows[i]
    row.cells[0].text = b
    row.cells[1].text = t
    row.cells[2].text = n
    if i >= 6:
        for c in row.cells:
            for p in c.paragraphs:
                for r in p.runs:
                    r.bold = True
                    r.font.color.rgb = RGBColor(31, 73, 125)

d.add_paragraph()

add_heading('2. Tiêu chí PICOS để chọn 22 bài vào SR', level=1)
add_para('Mỗi RCT phải thoả CẢ 5 điều kiện dưới đây. Tác giả công bố tiêu chí này TRƯỚC khi tìm kiếm (protocol PROSPERO CRD42023424181) để tránh bias lựa chọn.')

d.add_paragraph()

picos_tbl = d.add_table(rows=6, cols=3)
picos_tbl.style = 'Table Grid'

picos_hdr = picos_tbl.rows[0]
for i, h in enumerate(['Chữ', 'Yếu tố', 'Tiêu chí của Walder']):
    cell = picos_hdr.cells[i]
    shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h)
    r.bold = True
    r.font.color.rgb = RGBColor(255, 255, 255)

picos_data = [
    ('P', 'Population (Đối tượng)', 'Tuổi trung bình 0 đến <25 (trẻ em, vị thành niên, thanh niên)'),
    ('I', 'Intervention (Can thiệp)', 'DMHI (Digital Mental Health Intervention) dựa trên nguyên lý tâm lý, sử dụng từ xa (internet, app, VR)'),
    ('C', 'Comparator (Đối chứng)', 'Active intervention / waitlist (chờ) / care-as-usual (CAU) / treatment-as-usual (TAU)'),
    ('O', 'Outcome (Kết quả)', 'Có thang đo lo âu xã hội (social anxiety)'),
    ('S', 'Study design (Thiết kế)', 'RCT, đăng trên tạp chí peer-reviewed, ngôn ngữ tiếng Anh hoặc tiếng Đức'),
]
for i, (k, name, crit) in enumerate(picos_data, start=1):
    row = picos_tbl.rows[i]
    cell0 = row.cells[0]
    shade_cell(cell0, 'FFE699')
    p0 = cell0.paragraphs[0]
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r0 = p0.add_run(k)
    r0.bold = True
    r0.font.size = Pt(14)
    r0.font.color.rgb = RGBColor(191, 97, 14)
    row.cells[1].text = name
    for p in row.cells[1].paragraphs:
        for r in p.runs:
            r.bold = True
    row.cells[2].text = crit

d.add_paragraph()

excl_tbl = d.add_table(rows=1, cols=1)
excl_tbl.style = 'Table Grid'
ec = excl_tbl.rows[0].cells[0]
shade_cell(ec, 'FCE4D6')
ep = ec.paragraphs[0]
er = ep.add_run('LOẠI TRỪ: ')
er.bold = True
er.font.color.rgb = RGBColor(192, 0, 0)
ep.add_run('study protocols, review bài, các meta-analysis khác.')

d.add_paragraph()

add_heading('3. Tại sao 22 → 21 (chỉ có 21 bài vào MA)?', level=1)
add_para('1 bài duy nhất bị loại khỏi meta-analysis là Vigerland et al. [ref 93] — lý do trích nguyên văn từ bài gốc:')

d.add_paragraph()

quote_tbl = d.add_table(rows=1, cols=1)
quote_tbl.style = 'Table Grid'
qc = quote_tbl.rows[0].cells[0]
shade_cell(qc, 'DEEBF7')
qp = qc.paragraphs[0]
qr = qp.add_run('"The study by Vigerland and colleagues [93] was not included in the meta-analytic calculations because there was insufficient data of social anxiety at post-assessment."')
qr.italic = True
qr.font.color.rgb = RGBColor(31, 73, 125)
qp2 = qc.add_paragraph()
qr2 = qp2.add_run('(Tạm dịch: Bài của Vigerland và cộng sự [93] không được đưa vào tính toán meta-analysis vì thiếu dữ liệu lo âu xã hội tại thời điểm hậu can thiệp.)')
qr2.font.size = Pt(10)
qr2.font.color.rgb = RGBColor(90, 90, 90)

d.add_paragraph()
add_para('→ Bài này VẪN NẰM TRONG SR (được mô tả định tính trong bảng Study Characteristics), nhưng KHÔNG gộp vào meta-analysis vì không có con số (mean, SD, n) tại post-test để tính effect size Hedges g.')

d.add_paragraph()

add_heading('4. Ghi nhớ nhanh', level=1)

memo_tbl = d.add_table(rows=3, cols=1)
memo_tbl.style = 'Table Grid'

cells_content = [
    ('SR = 22 bài', 'Mô tả tổng quan định tính (ai, đâu, làm gì, đo gì)', 'E2EFDA'),
    ('MA = 21 bài', 'Chỉ những bài có ĐỦ số liệu post-assessment để tính effect size g', 'E2EFDA'),
    ('Lý do SR > MA', 'THIẾU DỮ LIỆU OUTCOME tại thời điểm post-test — đây là lý do thường gặp nhất khi số bài SR nhiều hơn MA 1-2 bài', 'FFF2CC'),
]
for i, (head, body, color) in enumerate(cells_content):
    cell = memo_tbl.rows[i].cells[0]
    shade_cell(cell, color)
    p = cell.paragraphs[0]
    r1 = p.add_run(head + ': ')
    r1.bold = True
    r1.font.size = Pt(12)
    r1.font.color.rgb = RGBColor(31, 73, 125)
    p.add_run(body)

d.add_paragraph()

add_heading('5. Bài học mở rộng — Khi nào cần lưu ý SR ≠ MA?', level=1, color=(192, 80, 77))

lessons = [
    ('Chênh lệch 1-2 bài (SR > MA)', 'BÌNH THƯỜNG. Thường do 1-2 bài thiếu số liệu post-test. Walder 2025 (22 vs 21) là ví dụ điển hình.'),
    ('Chênh lệch ≥ 5 bài', 'CẦN ĐỌC KỸ PRISMA flow. Có thể reviewer loại bỏ thêm theo tiêu chí methodological (risk of bias cao, heterogeneity, outlier).'),
    ('SR không có MA', 'Xảy ra khi heterogeneity quá cao (I² > 90%) hoặc các bài dùng thang đo quá khác nhau không pool được. Khi đó chỉ tổng quan định tính.'),
    ('MA > SR (nghịch lý)', 'KHÔNG THỂ. MA luôn là tập con của SR. Nếu thấy paper ghi ngược lại → lỗi báo cáo hoặc nghi ngờ methodology.'),
]

for situ, expl in lessons:
    p = d.add_paragraph()
    r1 = p.add_run('• ' + situ + ': ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(192, 80, 77)
    p.add_run(expl)

d.add_paragraph()

add_heading('Tài liệu tham khảo', level=1)
refs = [
    'Walder N, Frey A, Berger T, et al. (2025). Digital mental health interventions for prevention and treatment of social anxiety disorder in children, adolescents and young adults: Systematic review and meta-analysis. Journal of Medical Internet Research. DOI: 10.2196/preprints.67067.',
    'Page MJ, McKenzie JE, Bossuyt PM, et al. (2021). The PRISMA 2020 statement: an updated guideline for reporting systematic reviews. BMJ, 372:n71.',
    'Richardson WS, Wilson MC, Nishikawa J, Hayward RS (1995). The well-built clinical question: a key to evidence-based decisions. ACP Journal Club, 123(3):A12-3. [PICOS framework]',
    'Cochrane Handbook for Systematic Reviews of Interventions, v6.4 (2023), Chapter 6.',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs:
        r.font.size = Pt(10)

d.add_paragraph()
meta_foot = d.add_paragraph()
mf = meta_foot.add_run('Biên soạn: 20/04/2026 | Mọi số liệu đã verify trực tiếp từ PDF QT040 bản gốc.')
mf.font.size = Pt(9)
mf.italic = True
mf.font.color.rgb = RGBColor(128, 128, 128)

out = '01_Bao-cao/QT040_Walder_2025_Quy_trinh_loc_22_vs_21.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
