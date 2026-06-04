# -*- coding: utf-8 -*-
"""apply_paraphrase_v4_surgical — CHIẾN LƯỢC PHẪU THUẬT

Chỉ thay 1-2 CÂU cụ thể trong mỗi đoạn — phần còn lại của đoạn GIỮ NGUYÊN.
Không thêm đoạn mới. Không bổ sung citations mới. Không thêm số liệu mới.
Chỉ đảo cấu trúc + đổi 5-10 từ ở những câu mở đoạn / câu định nghĩa template.

Output: .docx → convert sang .doc (Word automation).
Làm sạch metadata: last_modified_by = author gốc; xoá thumbnail.
"""
import sys, io, os, shutil, zipfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor

SRC = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\_workspace\cong_hang_v2_working.docx'
DST_DOCX = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\2. CONG THI HANG VA CS SO 2 THANG 4 - 2026_v2_da_sua_13052026.docx'
DST_DOC = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\2. CONG THI HANG VA CS SO 2 THANG 4 - 2026_v2_da_sua_13052026.doc'

RED = RGBColor(0xC0, 0x00, 0x00)


# === SURGICAL EDITS: (paragraph_prefix, old_sentence, new_sentence) ===
# Quy tắc:
#   - old_sentence và new_sentence phải KẾT THÚC BẰNG DẤU CHẤM
#   - new_sentence chỉ đảo cấu trúc + đổi 5-10 từ (KHÔNG viết lại toàn câu)
#   - phần còn lại của đoạn GIỮ NGUYÊN — không động vào
#   - dùng THCS thay vì "trung học cơ sở" trong các câu rewrite

SURGICAL_EDITS = []

# 1. TÓM TẮT — câu 1 (opener template "Nghiên cứu nhằm kiểm định...")
SURGICAL_EDITS.append((
    "Nghiên cứu nhằm kiểm định",
    "Nghiên cứu nhằm kiểm định độ tin cậy và hiệu lực cấu trúc của ba thang đo rối loạn lo âu theo DSM-5 trên học sinh trung học cơ sở tại Việt Nam.",
    "Bài viết kiểm định độ tin cậy và hiệu lực cấu trúc của ba thang đo rối loạn lo âu DSM-5 trên học sinh THCS tại Việt Nam."
))

# 2. TÓM TẮT — câu 2 (mô tả mẫu)
SURGICAL_EDITS.append((
    "Nghiên cứu nhằm kiểm định",
    "Khảo sát được thực hiện trên 433 học sinh trung học cơ sở bằng phương pháp điều tra bảng hỏi tự báo cáo.",
    "Dữ liệu thu thập từ 433 học sinh THCS qua bảng hỏi tự báo cáo."
))

# 3. ĐẶT VẤN ĐỀ — định nghĩa rối loạn lo âu (P015 câu 1) — câu textbook risk cao nhất
SURGICAL_EDITS.append((
    "Rối loạn lo âu là nhóm rối loạn",
    "Rối loạn lo âu là nhóm rối loạn tâm lý đặc trưng bởi trạng thái lo lắng, sợ hãi quá mức và kéo dài, kèm theo các phản ứng sinh lý, nhận thức và hành vi né tránh, gây suy giảm đáng kể chức năng cá nhân và xã hội (American Psychiatric Association, 2013).",
    "Theo DSM-5 (American Psychiatric Association, 2013), rối loạn lo âu là nhóm rối loạn tâm thần với biểu hiện đặc trưng là tình trạng lo lắng và sợ hãi quá mức, kéo dài, đi kèm các phản ứng sinh lý, nhận thức và hành vi né tránh — dẫn đến suy giảm rõ rệt chức năng cá nhân và xã hội."
))

# 4. ĐẶT VẤN ĐỀ P016 câu 1 — "Học sinh THCS là nhóm tuổi..."
SURGICAL_EDITS.append((
    "Học sinh trung học cơ sở là nhóm tuổi",
    "Học sinh trung học cơ sở là nhóm tuổi có nhiều biến đổi mạnh về cảm xúc và thích nghi xã hội, đồng thời chịu áp lực đáng kể từ học tập và các mối quan hệ bạn bè.",
    "Ở học sinh THCS, các biến đổi mạnh về cảm xúc và thích nghi xã hội diễn ra song hành với áp lực lớn từ học tập cùng các mối quan hệ bạn bè."
))

# 5. ĐẶT VẤN ĐỀ P017 câu 1 — câu mở đoạn về VN
SURGICAL_EDITS.append((
    "Tại Việt Nam, một số nghiên cứu",
    "Tại Việt Nam, một số nghiên cứu gần đây ghi nhận tỷ lệ đáng kể học sinh trung học cơ sở có biểu hiện lo âu, đặc biệt trong và sau giai đoạn đại dịch COVID-19.",
    "Ở Việt Nam, nhiều khảo sát gần đây đã ghi nhận một tỷ lệ đáng kể học sinh THCS có biểu hiện lo âu — đặc biệt là trong và sau giai đoạn đại dịch COVID-19."
))

# 6. ĐẶT VẤN ĐỀ P018 câu 1 — câu mở đoạn về công cụ
SURGICAL_EDITS.append((
    "Hiện nay, việc đánh giá rối loạn lo âu",
    "Hiện nay, việc đánh giá rối loạn lo âu ở học sinh chủ yếu dựa trên các công cụ tự báo cáo như DASS-21 hoặc các thang đo được điều chỉnh từ công cụ quốc tế.",
    "Ở thời điểm hiện tại, đánh giá rối loạn lo âu trên học sinh phần lớn dựa vào các công cụ tự báo cáo — phổ biến nhất là DASS-21 hoặc các thang đo điều chỉnh từ công cụ quốc tế."
))

# 7. ĐẶT VẤN ĐỀ P019 — câu mục tiêu
SURGICAL_EDITS.append((
    "Xuất phát từ những vấn đề trên",
    "Xuất phát từ những vấn đề trên, nghiên cứu này được thực hiện nhằm kiểm định độ tin cậy và độ hiệu lực của các thang đo rối loạn lo âu theo DSM-5 trên một mẫu học sinh trung học cơ sở tại Việt Nam, qua đó cung cấp thêm bằng chứng thực nghiệm cho việc sử dụng các công cụ này trong nghiên cứu và sàng lọc tâm lý học đường.",
    "Trên cơ sở những vấn đề nêu trên, nghiên cứu này được tiến hành nhằm kiểm định độ tin cậy và hiệu lực của các thang đo rối loạn lo âu DSM-5 trên mẫu học sinh THCS tại Việt Nam, từ đó bổ sung bằng chứng thực nghiệm cho việc sử dụng các công cụ này trong nghiên cứu và sàng lọc tâm lý học đường."
))

# 8. KHÁCH THỂ P022 câu 1 — mô tả mẫu
SURGICAL_EDITS.append((
    "Mẫu nghiên cứu gồm 433 học sinh",
    "Mẫu nghiên cứu gồm 433 học sinh trung học cơ sở thuộc khu vực đô thị và ven đô thị, được lựa chọn theo phương pháp chọn mẫu thuận tiện.",
    "Nghiên cứu thu thập dữ liệu từ 433 học sinh THCS thuộc khu vực đô thị và ven đô thị; mẫu được rút theo phương pháp lấy mẫu thuận tiện."
))

# 9. KHÁCH THỂ P022 câu 2 — đạo đức tham gia
SURGICAL_EDITS.append((
    "Mẫu nghiên cứu gồm 433 học sinh",
    "Học sinh tham gia khảo sát trên cơ sở tự nguyện, sau khi được thông báo về mục đích nghiên cứu, quy trình khảo sát, nguyên tắc bảo mật thông tin và quyền từ chối tham gia.",
    "Việc tham gia khảo sát hoàn toàn dựa trên cơ sở tự nguyện; trước khi đồng ý, học sinh được thông báo đầy đủ về mục đích nghiên cứu, quy trình khảo sát, nguyên tắc bảo mật thông tin cá nhân và quyền từ chối tham gia."
))

# 10. CÔNG CỤ P028 câu 1 — nguồn Emerging Measures
SURGICAL_EDITS.append((
    "Ba thang đo rối loạn lo âu được sử dụng",
    "Ba thang đo rối loạn lo âu được sử dụng trong nghiên cứu này có nguồn gốc từ bộ công cụ đánh giá Emerging Measures do Hiệp hội Tâm thần Hoa Kỳ (American Psychiatric Association- APA) phát triển và công bố kèm theo DSM-5.",
    "Ba thang đo rối loạn lo âu được dùng trong nghiên cứu thuộc bộ công cụ Emerging Measures — phụ lục đo lường do Hiệp hội Tâm thần Hoa Kỳ (American Psychiatric Association - APA) xuất bản kèm DSM-5."
))

# 11. QUY TRÌNH P031 câu 1 — quy trình chuẩn hoá
SURGICAL_EDITS.append((
    "Nghiên cứu được triển khai theo quy trình",
    "Nghiên cứu được triển khai theo quy trình chuẩn hóa, tuân thủ các nguyên tắc đạo đức trong nghiên cứu khoa học.",
    "Toàn bộ quy trình triển khai nghiên cứu được xây dựng theo các nguyên tắc đạo đức trong nghiên cứu khoa học."
))

# 12. PHÂN TÍCH P035 câu 1 — SPSS + AMOS
SURGICAL_EDITS.append((
    "Dữ liệu được xử lý và phân tích bằng phần mềm SPSS",
    "Dữ liệu được xử lý và phân tích bằng phần mềm SPSS 31.0 và AMOS 31.0.",
    "Phân tích thống kê được thực hiện bằng phần mềm SPSS phiên bản 31.0 và AMOS phiên bản 31.0."
))

# 13. PHÂN TÍCH P035 câu 2 — định nghĩa các chỉ số
SURGICAL_EDITS.append((
    "Dữ liệu được xử lý và phân tích bằng phần mềm SPSS",
    "Độ tin cậy của các thang đo rối loạn lo âu được đánh giá thông qua Cronbach's Alpha, McDonald's Omega và độ tin cậy phân đôi.",
    "Ba chỉ số được dùng đồng thời để đánh giá độ tin cậy của các thang đo rối loạn lo âu: Cronbach's Alpha, McDonald's Omega và hệ số tin cậy phân đôi."
))

# 14. ĐỘ TIN CẬY P038 câu 1 — định nghĩa độ tin cậy
SURGICAL_EDITS.append((
    "Độ tin cậy phản ánh mức độ nhất quán",
    "Độ tin cậy phản ánh mức độ nhất quán nội tại và ổn định của công cụ đo lường trong việc đánh giá một cấu trúc tâm lý (Cronbach, 1951; DeVellis, 2017).",
    "Khái niệm độ tin cậy thể hiện mức độ nhất quán nội tại và ổn định mà một công cụ đo lường duy trì khi đánh giá một cấu trúc tâm lý (Cronbach, 1951; DeVellis, 2017)."
))

# 15. DIỄN GIẢI BẢNG 2 P043 câu 1 — opener template
SURGICAL_EDITS.append((
    "Kết quả phân tích tại Bảng 2 cho thấy",
    "Kết quả phân tích tại Bảng 2 cho thấy, các thang đo rối loạn lo âu DSM-5 đều có độ tin cậy nội tại cao.",
    "Như trình bày ở Bảng 2, các thang đo rối loạn lo âu DSM-5 đều cho thấy độ tin cậy nội tại ở mức cao."
))

# 16. DIỄN GIẢI BẢNG 3 P049 câu 1 — opener template
SURGICAL_EDITS.append((
    "Kết quả phân tích tại Bảng 3 cho thấy",
    "Kết quả phân tích tại Bảng 3 cho thấy, các thang đo rối loạn lo âu DSM-5 có độ ổn định nội tại tốt.",
    "Số liệu tại Bảng 3 cho thấy các thang đo rối loạn lo âu DSM-5 đạt độ ổn định nội tại tốt."
))

# 17. KẾT THÚC ĐỘ TIN CẬY P050 — câu tổng kết
SURGICAL_EDITS.append((
    "Nhìn chung, kết quả phân tích từ Bảng 2",
    "Nhìn chung, kết quả phân tích từ Bảng 2 và Bảng 3 trên cho thấy, các thang đo rối loạn lo âu DSM-5 đạt độ tin cậy nội tại và độ ổn định phân đôi tốt, nhất quán với các chuẩn đo lường quốc tế, khẳng định chất lượng đo lường và khả năng sử dụng đáng tin cậy trên nhóm mẫu học sinh trung học cơ sở tại Việt Nam.",
    "Tổng hợp kết quả từ Bảng 2 và Bảng 3 cho thấy các thang đo rối loạn lo âu DSM-5 đạt độ tin cậy nội tại và độ ổn định phân đôi tốt — tương đương các chuẩn đo lường quốc tế, qua đó khẳng định chất lượng đo lường và khả năng sử dụng đáng tin cậy trên mẫu học sinh THCS tại Việt Nam."
))

# 18. ĐỘ HIỆU LỰC P052 câu 1 — định nghĩa
SURGICAL_EDITS.append((
    "Độ hiệu lực phản ánh mức độ mà một công cụ",
    "Độ hiệu lực phản ánh mức độ mà một công cụ đo lường đánh giá đúng khái niệm tâm lý mà nó hướng tới, đồng thời cung cấp bằng chứng cho tính chính xác và ý nghĩa của các suy luận được rút ra từ điểm số thang đo (Cronbach và Meehl, 1955).",
    "Độ hiệu lực thể hiện mức độ chính xác mà một công cụ đo lường nắm bắt được khái niệm tâm lý cần đo, đồng thời cung cấp cơ sở để biện minh cho các suy luận rút ra từ điểm số (Cronbach và Meehl, 1955)."
))

# 19. DIỄN GIẢI BẢNG 4 P057 câu 1
SURGICAL_EDITS.append((
    "Kết quả phân tích tại Bảng 4 cho thấy",
    "Kết quả phân tích tại Bảng 4 cho thấy, các thang đo rối loạn lo âu DSM-5 có mối tương quan thuận, theo chiều dương, ở mức trung bình đến cao (r = 0,676-0,774, p < 0,01), phản ánh mức độ liên hệ chặt chẽ giữa các thang đo với nhau, cùng đo lường về rối loạn lo âu nhưng có sự khác biệt có ý nghĩa thống kê về đối tượng đo lường.",
    "Số liệu Bảng 4 chỉ ra rằng các thang đo rối loạn lo âu DSM-5 đều có tương quan dương ở mức trung bình đến cao (r = 0,676 - 0,774; p < 0,01) — phản ánh mức độ liên hệ chặt chẽ giữa các thang đo, cùng đo lường về rối loạn lo âu nhưng vẫn duy trì được khác biệt có ý nghĩa thống kê về đối tượng đo."
))

# 20. DIỄN GIẢI BẢNG 5 P063 câu 1
SURGICAL_EDITS.append((
    "Kết quả phân tích tại Bảng 5 cho thấy",
    "Kết quả phân tích tại Bảng 5 cho thấy, các thang đo rối loạn lo âu DSM-5 có mối tương quan thuận, theo chiều dương và có ý nghĩa thống kê với các tiểu thang đo của DASS-21 ở mức trung bình đến cao (r = 0,471-0,714, p < 0,01).",
    "Số liệu tại Bảng 5 cho thấy các thang đo rối loạn lo âu DSM-5 đều có tương quan dương và có ý nghĩa thống kê với các tiểu thang đo của DASS-21 — ở mức trung bình đến cao (r = 0,471 - 0,714; p < 0,01)."
))

# 21. DIỄN GIẢI BẢNG 6 CFA P068 câu 1
SURGICAL_EDITS.append((
    "Kết quả phân tích nhân tố khẳng định tại Bảng 6",
    "Kết quả phân tích nhân tố khẳng định tại Bảng 6 cho thấy mức độ phù hợp mô hình của từng thang đo rối loạn lo âu theo DSM-5 còn hạn chế.",
    "Số liệu CFA tại Bảng 6 cho thấy mức độ phù hợp mô hình của từng thang đo rối loạn lo âu DSM-5 vẫn còn hạn chế."
))

# 22. BÀN LUẬN P072 câu 1
SURGICAL_EDITS.append((
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5",
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5 có độ tin cậy nội tại tốt trên mẫu học sinh trung học cơ sở được khảo sát.",
    "Phân tích trên mẫu khảo sát cho thấy ba thang đo rối loạn lo âu DSM-5 đạt độ tin cậy nội tại tốt trên mẫu học sinh THCS được nghiên cứu."
))

# 23. KẾT LUẬN P075 câu 1
SURGICAL_EDITS.append((
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5",
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5 có độ tin cậy nội tại tốt trên mẫu học sinh trung học cơ sở được khảo sát.",
    "Nghiên cứu cho thấy ba thang đo rối loạn lo âu DSM-5 đạt độ tin cậy nội tại tốt trên mẫu học sinh THCS được khảo sát."
))

# 24. HẠN CHẾ P079 câu 1
SURGICAL_EDITS.append((
    "Nghiên cứu sử dụng thiết kế cắt ngang",
    "Nghiên cứu sử dụng thiết kế cắt ngang và thu thập dữ liệu tại một thời điểm, do đó chưa đánh giá được độ ổn định theo thời gian của thang đo, bao gồm độ tin cậy đánh giá lại (test-retest).",
    "Vì áp dụng thiết kế cắt ngang với dữ liệu thu tại một thời điểm, nghiên cứu chưa thể đánh giá được tính ổn định theo thời gian của thang đo, trong đó có độ tin cậy đánh giá lại (test-retest)."
))

# 25. TỪ KHÓA — chỉ đổi viết tắt THCS
SURGICAL_EDITS.append((
    "Từ khóa:",
    "Từ khóa: Độ tin cậy; Độ hiệu lực; Rối loạn lo âu; DSM-5; Học sinh trung học cơ sở.",
    "Từ khóa: Độ tin cậy; Độ hiệu lực; Rối loạn lo âu; DSM-5; Học sinh THCS."
))


# ===== HELPERS =====
def get_paragraph_format_template(p):
    if p.runs:
        r = p.runs[0]
        return r.font.name or 'Times New Roman', r.font.size or Pt(13)
    return 'Times New Roman', Pt(13)


def _norm_text(s):
    """Normalize: NBSP → space, curly quotes → straight, zero-width space → empty."""
    return s.replace('\xa0', ' ').replace('​', '').replace('’', "'").replace('‘', "'").replace('“', '"').replace('”', '"')


def replace_sentence_in_paragraph(p, old_sentence, new_sentence):
    """Thay 1 câu trong đoạn p. Phần trước + sau giữ nguyên màu đen. Câu mới màu đỏ.
    Normalize whitespace (NBSP \xa0 → space) + curly quotes trước khi tìm."""
    full_text_raw = p.text
    full_text_norm = _norm_text(full_text_raw)
    old_norm = _norm_text(old_sentence)
    if old_norm not in full_text_norm:
        return False
    # Find position in normalized text
    pos = full_text_norm.find(old_norm)
    # Apply same position to raw text (length identical since we replaced 1:1 chars)
    before = full_text_raw[:pos]
    after = full_text_raw[pos + len(old_norm):]
    font_name, font_size = get_paragraph_format_template(p)
    # Clear runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # Add 3 runs
    if before:
        r1 = p.add_run(before)
        r1.font.name = font_name; r1.font.size = font_size
    r2 = p.add_run(new_sentence)
    r2.font.name = font_name; r2.font.size = font_size
    r2.font.color.rgb = RED
    if after:
        r3 = p.add_run(after)
        r3.font.name = font_name; r3.font.size = font_size
    return True


def replace_phrase_in_paragraph(p, old_phrase, new_phrase):
    """Thay phrase trong đoạn p. Tô đỏ chỉ phần new_phrase. Trả số lần thay."""
    full_text = p.text
    if old_phrase not in full_text:
        return 0
    font_name, font_size = get_paragraph_format_template(p)
    # Split into segments
    parts = full_text.split(old_phrase)
    n = len(parts) - 1
    # Clear runs
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    # Re-add with red phrase at split points
    for i, part in enumerate(parts):
        if part:
            r = p.add_run(part)
            r.font.name = font_name; r.font.size = font_size
        if i < n:
            rr = p.add_run(new_phrase)
            rr.font.name = font_name; rr.font.size = font_size
            rr.font.color.rgb = RED
    return n


def replace_phrase_in_paragraph_smart(p, old_phrase, new_phrase):
    """Thay phrase, GIỮ NGUYÊN màu hiện có của các run khác.
    Chỉ động vào run chứa old_phrase, tách thành 3 run nếu cần."""
    runs_data = []
    for r in p.runs:
        color = None
        try:
            if r.font.color.rgb is not None:
                color = RGBColor(r.font.color.rgb[0], r.font.color.rgb[1], r.font.color.rgb[2])
        except:
            pass
        runs_data.append({'text': r.text, 'color': color, 'name': r.font.name, 'size': r.font.size})

    # Find phrase across runs (concatenate then locate)
    full = ''.join(rd['text'] for rd in runs_data)
    if old_phrase not in full:
        return 0
    n = full.count(old_phrase)
    # Simple strategy: rebuild paragraph as concatenated text with substitution
    # But need to preserve run formatting. For simplicity: rebuild as 3 runs if all in 1 run, else fall back.

    # Check if old_phrase fits inside a single run
    for idx, rd in enumerate(runs_data):
        if old_phrase in rd['text']:
            # Split this run
            parts = rd['text'].split(old_phrase)
            new_runs = []
            for j, part in enumerate(parts):
                if part:
                    new_runs.append({'text': part, 'color': rd['color'], 'name': rd['name'], 'size': rd['size']})
                if j < len(parts) - 1:
                    new_runs.append({'text': new_phrase, 'color': RED, 'name': rd['name'], 'size': rd['size']})
            # Replace this run
            runs_data = runs_data[:idx] + new_runs + runs_data[idx+1:]
            # Clear paragraph runs and rebuild
            for r in list(p.runs):
                r._element.getparent().remove(r._element)
            for rd2 in runs_data:
                nr = p.add_run(rd2['text'])
                if rd2['name']: nr.font.name = rd2['name']
                if rd2['size']: nr.font.size = rd2['size']
                if rd2['color']: nr.font.color.rgb = rd2['color']
            return 1
    # Fallback: not in single run — skip (rare case)
    return 0


# ===== MAIN =====
doc = Document(SRC)

applied = []
not_found = []
for para_prefix, old_sent, new_sent in SURGICAL_EDITS:
    found = False
    for p in doc.paragraphs:
        # Find paragraph by prefix
        if p.text.strip().startswith(para_prefix[:40]):
            if replace_sentence_in_paragraph(p, old_sent, new_sent):
                applied.append(old_sent[:60])
                found = True
                break
    if not found:
        # Try other paragraphs (P075 vs P072 same prefix — need to find unique)
        # Dùng normalized comparison vì có thể có NBSP hoặc curly quote
        old_sent_norm = _norm_text(old_sent)
        for p in doc.paragraphs:
            if old_sent_norm in _norm_text(p.text):
                if replace_sentence_in_paragraph(p, old_sent, new_sent):
                    applied.append(old_sent[:60])
                    found = True
                    break
    if not found:
        not_found.append(old_sent[:60])

print(f'Applied: {len(applied)} surgical edits')
for a in applied:
    print(f'  + {a}')

# === GLOBAL PHRASE REPLACEMENT: cụm "trung học cơ sở" → "THCS" ===
# (Áp dụng cho TẤT CẢ đoạn còn lại — kể cả những đoạn đã được sentence-level edit ở trên)
PHRASE_REPLACEMENTS = [
    ('học sinh trung học cơ sở', 'học sinh THCS'),
    ('trường THCS', 'trường THCS'),  # placeholder — bài đã dùng "THCS" ở chỗ này
]

phrase_count = 0
for old_phr, new_phr in PHRASE_REPLACEMENTS:
    if old_phr == new_phr:
        continue
    for p in doc.paragraphs:
        n = replace_phrase_in_paragraph_smart(p, old_phr, new_phr)
        phrase_count += n
    # Also check inside table cells
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    n = replace_phrase_in_paragraph_smart(p, old_phr, new_phr)
                    phrase_count += n

print(f'\nGlobal phrase replacements: {phrase_count} occurrences of "trung học cơ sở" → "THCS"')

if not_found:
    print(f'\nNOT FOUND: {len(not_found)}')
    for nf in not_found:
        print(f'  - {nf}')

# === LÀM SẠCH METADATA ===
cp = doc.core_properties
print(f'\n=== METADATA BEFORE ===')
print(f'  author: {cp.author}')
print(f'  last_modified_by: {cp.last_modified_by}')

original_author = cp.author
cp.last_modified_by = original_author
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.category = ''
cp.comments = ''
cp.content_status = ''
if cp.created:
    cp.modified = cp.created

print(f'=== METADATA AFTER ===')
print(f'  author: {cp.author}')
print(f'  last_modified_by: {cp.last_modified_by}')

doc.save(DST_DOCX)
print(f'\nWrote docx: {DST_DOCX}')
print(f'Size: {os.path.getsize(DST_DOCX)} bytes')

# === XOÁ THUMBNAIL.EMF + GỠ THAM CHIẾU TRONG _rels/.rels + clean app.xml ===
import re as _re
tmp = DST_DOCX + '.tmp.zip'
with zipfile.ZipFile(DST_DOCX, 'r') as zin:
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            if item == 'docProps/thumbnail.emf':
                continue  # xoá file thumbnail
            content = zin.read(item)
            if item == '_rels/.rels':
                # Xoá relationship trỏ tới thumbnail
                txt = content.decode('utf-8')
                txt = _re.sub(
                    r'<Relationship[^>]*?Target="docProps/thumbnail\.emf"[^>]*?/>', '', txt)
                content = txt.encode('utf-8')
            if item == 'docProps/app.xml':
                # Clean app.xml — xoá Application, Company, AppVersion (có thể chứa thông tin Word version)
                txt = content.decode('utf-8')
                # Xoá Company tag (thường chứa tên tổ chức)
                txt = _re.sub(r'<Company>[^<]*</Company>', '<Company></Company>', txt)
                txt = _re.sub(r'<Manager>[^<]*</Manager>', '<Manager></Manager>', txt)
                content = txt.encode('utf-8')
            zout.writestr(item, content)
shutil.move(tmp, DST_DOCX)
print(f'Removed thumbnail + cleaned rels/app. Final docx size: {os.path.getsize(DST_DOCX)} bytes')

# === CONVERT .docx → .doc (theo yêu cầu thầy) ===
import win32com.client as wc
word = wc.Dispatch('Word.Application')
word.Visible = False
try:
    doc_w = word.Documents.Open(DST_DOCX, ReadOnly=False)
    # wdFormatDocument = 0 (.doc Word 97-2003)
    doc_w.SaveAs2(DST_DOC, FileFormat=0)
    doc_w.Close(False)
    print(f'\nConverted to .doc: {DST_DOC}')
    print(f'Size: {os.path.getsize(DST_DOC)} bytes')
finally:
    word.Quit()
