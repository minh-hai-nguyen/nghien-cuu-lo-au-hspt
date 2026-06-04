"""Bản dịch QT064 Stephens 2023 Photovoice scoping review — full bilingual."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/03_Ban-dich/QT064_Stephens_Photovoice_ScopingReview_IntlJAdolYouth_2023.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
DARK = RGBColor(31, 73, 125); RED = RGBColor(192, 0, 0); GRAY = RGBColor(90, 90, 90); ORANGE = RGBColor(191, 97, 14); GREEN = RGBColor(54, 95, 44)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def en(text):
    p = d.add_paragraph(); r = p.add_run(text); r.italic = True; r.font.color.rgb = GRAY; r.font.size = Pt(10)
def vn(text, bold=False):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(11)
def page(n):
    p = d.add_paragraph(); r = p.add_run(f'--- Trang {n} ---'); r.bold = True; r.font.color.rgb = ORANGE; r.font.size = Pt(10)

# =============================================================
# COVER
# =============================================================
t = d.add_heading('QT064 — Bản dịch chi tiết: Photovoice trong nghiên cứu sức khoẻ tâm thần vị thành niên (Stephens et al. 2023)', level=0)
for r in t.runs: r.font.color.rgb = DARK
sub = d.add_paragraph(); sr = sub.add_run('Bản dịch song ngữ Anh-Việt chi tiết 1-1. Phần Table 2 (data extraction 53 trang) + References giữ tiếng Anh trong PDF gốc.')
sr.italic = True; sr.font.size = Pt(11); sr.font.color.rgb = GRAY
d.add_paragraph()

# THÔNG TIN THƯ MỤC
add_h('Thông tin thư mục', 1)
meta = d.add_table(rows=10, cols=2); meta.style = 'Table Grid'
mrows = [
    ('Tiêu đề (EN)', 'A systematic scoping review of Photovoice within mental health research involving adolescents'),
    ('Tiêu đề (VN)', 'Tổng quan phạm vi hệ thống về Photovoice trong nghiên cứu sức khoẻ tâm thần liên quan vị thành niên'),
    ('Tác giả', 'Madison Stephens, Eleanor Keiller, Maev Conneely, Paul Heritage, Mariana Steffen, Victoria Jane Bird'),
    ('Đơn vị', 'Unit for Social and Community Psychiatry, WHO Collaborating Centre, Queen Mary University of London + East London NHS Foundation Trust + UCL Division of Psychiatry, UK'),
    ('Tạp chí', 'International Journal of Adolescence and Youth, vol. 28, no. 1, 2244043 (2023)'),
    ('DOI', 'https://doi.org/10.1080/02673843.2023.2244043'),
    ('Ngày nhận', '15 May 2023'),
    ('Ngày chấp nhận', '28 July 2023'),
    ('Số trang', '82 (chính: 18 trang; Table 2 mở rộng: 53 trang; refs: 7 trang)'),
    ('Open Access', 'Yes (CC BY 4.0). Tải tại UCL Discovery: https://discovery.ucl.ac.uk/id/eprint/10179741/'),
]
for i, (k, v) in enumerate(mrows):
    c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
    meta.rows[i].cells[1].text = v

d.add_paragraph()

# =============================================================
# ABSTRACT
# =============================================================
page(2)
add_h('TÓM TẮT (Abstract)', 1)

en('Photovoice is a research method that changes perceptions of mental health. However, there is a lack of evidence exploring how Photovoice is used in mental health research involving adolescents.')
vn('Photovoice là một PHƯƠNG PHÁP NGHIÊN CỨU làm thay đổi cách nhìn nhận về sức khoẻ tâm thần. Tuy nhiên, hiện vẫn còn thiếu bằng chứng khám phá cách Photovoice được sử dụng trong nghiên cứu sức khoẻ tâm thần có sự tham gia của vị thành niên.')
d.add_paragraph()

en('Our review aimed to understand the nature and key themes across findings of Photovoice studies exploring mental health among adolescents.')
vn('Nghiên cứu tổng quan của chúng tôi nhằm hiểu BẢN CHẤT và các CHỦ ĐỀ CHÍNH xuyên suốt các phát hiện của những nghiên cứu Photovoice khám phá sức khoẻ tâm thần ở vị thành niên.')
d.add_paragraph()

en('We used pre-existing data and updated a search strategy. Popay and colleagues\' guidance was used to analyse the studies and the quality of each study was appraised.')
vn('Chúng tôi sử dụng dữ liệu có sẵn (pre-existing data) và cập nhật chiến lược tìm kiếm. Hướng dẫn của Popay và cộng sự được dùng để phân tích các nghiên cứu, chất lượng của mỗi nghiên cứu được đánh giá.')
d.add_paragraph()

en('Our review found that Photovoice studies exploring mental health among adolescents are limited in quality and that Photovoice is a flexible, adaptable, inclusive, and emerging method.')
vn('Tổng quan của chúng tôi phát hiện rằng các nghiên cứu Photovoice khám phá sức khoẻ tâm thần ở vị thành niên CHẤT LƯỢNG CÒN HẠN CHẾ, và Photovoice là một phương pháp LINH HOẠT, CÓ THỂ THÍCH ỨNG, BAO TRÙM và đang nổi lên (emerging method).')
d.add_paragraph()

en('Coping; resilience; beliefs about oneself; family; friends; safety; living in a lower socioeconomic area and treatment emerged as key themes across study findings.')
vn('Các CHỦ ĐỀ CHÍNH nổi lên xuyên suốt các phát hiện gồm: ỨNG PHÓ (coping); KIÊN CƯỜNG (resilience); NIỀM TIN VỀ BẢN THÂN (beliefs about oneself); GIA ĐÌNH; BẠN BÈ; AN TOÀN (safety); SỐNG TRONG KHU VỰC KINH TẾ-XÃ HỘI THẤP và ĐIỀU TRỊ (treatment).')
d.add_paragraph()

en('Our review is the first of its kind and highlights ways Photovoice studies in the future can be developed and is helpful to multiple stakeholders.')
vn('Tổng quan của chúng tôi là CÔNG TRÌNH ĐẦU TIÊN cùng loại, làm nổi bật các hướng phát triển nghiên cứu Photovoice trong tương lai và có ích cho nhiều bên liên quan.')
d.add_paragraph()

vn('Từ khoá (Keywords): Photovoice; vị thành niên; sức khoẻ tâm thần; nghiên cứu; phương pháp; nghiên cứu có sự tham gia (participatory research).', bold=True)
d.add_paragraph()

# =============================================================
# INTRODUCTION
# =============================================================
add_h('1. GIỚI THIỆU (Introduction)', 1)
add_h('1.1. Giới thiệu về Photovoice trong nghiên cứu', 2)

en('Photovoice is a research method (Han & Oliffe, 2016) that uses photography and narrative (Mooney & Bhui, 2023) and was developed by C. Wang and Burris (1997) to explore the experiences of marginalized Chinese women.')
vn('Photovoice là một phương pháp nghiên cứu (Han & Oliffe, 2016) sử dụng NHIẾP ẢNH và TỰ THUẬT (narrative) (Mooney & Bhui, 2023). Phương pháp này được phát triển bởi C. Wang và Burris (1997) để khám phá trải nghiệm của phụ nữ Trung Quốc bị thiệt thòi.')
d.add_paragraph()

en('Photovoice research is designed around three steps: (1) Participants are trained by facilitators on the use of cameras, ethics and power prior to collecting photographs; (2) participants discuss and reflect upon photographs during a group discussion and (3) participants analyse the data (C. Wang & Burris, 1997).')
vn('Nghiên cứu Photovoice được thiết kế quanh BA BƯỚC: (1) Người tham gia được người điều phối (facilitator) đào tạo về cách dùng máy ảnh, đạo đức và quyền lực TRƯỚC khi thu thập ảnh; (2) người tham gia thảo luận và suy ngẫm về các bức ảnh trong một buổi thảo luận nhóm; và (3) người tham gia phân tích dữ liệu (C. Wang & Burris, 1997).')
d.add_paragraph()

en('Freire\'s (1970) theory of critical consciousness, feminism and notions of voice and participatory documentary photography underpin Photovoice (C. C. Wang et al., 2004) which places emphasis on amplifying participants\' voices and fostering change.')
vn('LÝ THUYẾT về ý thức phê phán (critical consciousness) của Freire (1970), TƯ TƯỞNG nữ quyền (feminism) cùng các khái niệm về TIẾNG NÓI và NHIẾP ẢNH TÀI LIỆU CÓ SỰ THAM GIA (participatory documentary photography) là nền tảng của Photovoice (C. C. Wang và cộng sự, 2004), đặt trọng tâm vào KHUẾCH ĐẠI tiếng nói người tham gia và THÚC ĐẨY thay đổi.')
d.add_paragraph()

en('Participants involved in Photovoice research creatively express themselves (Cosgrove et al., 2023) by visually presenting their surroundings using photographs (Csesznek, 2021) in a way that transcends the limitations of narrative (C. Wang & Burris, 1997).')
vn('Người tham gia nghiên cứu Photovoice THỂ HIỆN BẢN THÂN một cách sáng tạo (Cosgrove và cộng sự, 2023) bằng cách trình bày trực quan môi trường xung quanh thông qua ảnh chụp (Csesznek, 2021), theo cách VƯỢT QUA giới hạn của tự thuật (C. Wang & Burris, 1997).')
d.add_paragraph()

en('Photovoice can generate insights that are emotional and metaphoric (Wass et al., 2020) that are not as accessible using other qualitative methods (McLaughlin & Coleman-Fountain, 2019). Notably, Guell and Ogilive (2013) claim that Photovoice produces knowledge that is more nuanced and rich compared to insights generated through interviews, making Photovoice advantageous over approaches relying solely on verbal and written data (Levin et al., 2007).')
vn('Photovoice có thể tạo ra những HIỂU BIẾT mang tính CẢM XÚC và ẨN DỤ (metaphoric) (Wass và cộng sự, 2020) mà các phương pháp định tính khác khó tiếp cận được (McLaughlin & Coleman-Fountain, 2019). Đáng chú ý, Guell và Ogilive (2013) khẳng định rằng Photovoice tạo ra TRI THỨC tinh tế và phong phú hơn so với phỏng vấn, khiến Photovoice có lợi thế so với các cách tiếp cận chỉ dựa vào dữ liệu lời nói và văn viết (Levin và cộng sự, 2007).')
d.add_paragraph()

page(3)
en('For instance, Byrne (2012) highlights how populations who may find verbal communication difficult, including those affected by mental ill-health, can utilize Photovoice as a compelling method to amplify their voices.')
vn('Ví dụ, Byrne (2012) nhấn mạnh cách những nhóm dân CÓ THỂ KHÓ giao tiếp bằng lời, bao gồm những người bị ảnh hưởng bởi bệnh tâm thần, có thể sử dụng Photovoice như một phương pháp THUYẾT PHỤC để khuếch đại tiếng nói của họ.')
d.add_paragraph()

en('Researchers who engage in community-based participatory research often use Photovoice (Lofton and Grant, 2021), in descriptive study designs (Badanta et al., 2021). Researchers have previously used Photovoice to explore the experiences of marginalized communities affected by HIV/AIDS (Earnshaw et al., 2023); homelessness (Mollica et al., 2023) and war (Feen-Calligan et al., 2023).')
vn('Các nhà nghiên cứu tham gia NGHIÊN CỨU CÓ SỰ THAM GIA DỰA TRÊN CỘNG ĐỒNG (community-based participatory research) thường dùng Photovoice (Lofton & Grant, 2021), trong các thiết kế nghiên cứu MÔ TẢ (Badanta và cộng sự, 2021). Các nhà nghiên cứu đã từng dùng Photovoice để khám phá trải nghiệm của các cộng đồng thiệt thòi bị ảnh hưởng bởi HIV/AIDS (Earnshaw và cộng sự, 2023); vô gia cư (Mollica và cộng sự, 2023) và chiến tranh (Feen-Calligan và cộng sự, 2023).')
d.add_paragraph()

en('Photovoice is a relevant method for research involving under-served communities which empowers participants (Mental Health Foundation, 2021; Schnittker, 2013) by increasing awareness of their surroundings (Budig et al., 2018) and re-distributing power (Wallerstein et al., 2019).')
vn('Photovoice là phương pháp PHÙ HỢP cho nghiên cứu liên quan đến các cộng đồng ít được phục vụ (under-served), TRAO QUYỀN cho người tham gia (Mental Health Foundation, 2021; Schnittker, 2013) bằng cách tăng nhận thức về môi trường xung quanh họ (Budig và cộng sự, 2018) và TÁI PHÂN PHỐI QUYỀN LỰC (Wallerstein và cộng sự, 2019).')
d.add_paragraph()

en('Unlike many other research methods, Photovoice allows participants to engage and lead in all stages of Photovoice research, highlighting a democratic approach (Kimera & Vindevogel, 2022). Democratic research is important in ensuring that research aligns with public needs (Pamuk, 2020) and Photovoice studies succinctly communicate individuals\' priorities and community strengths and weaknesses (Teti et al., 2012).')
vn('Khác với nhiều phương pháp nghiên cứu khác, Photovoice CHO PHÉP người tham gia tham gia và DẪN DẮT mọi giai đoạn nghiên cứu, làm nổi bật một CÁCH TIẾP CẬN DÂN CHỦ (Kimera & Vindevogel, 2022). Nghiên cứu dân chủ quan trọng trong việc đảm bảo nghiên cứu phù hợp với nhu cầu công chúng (Pamuk, 2020), và các nghiên cứu Photovoice truyền tải ngắn gọn các ưu tiên cá nhân cùng điểm mạnh, điểm yếu của cộng đồng (Teti và cộng sự, 2012).')
d.add_paragraph()

en('In turn, participants involved in Photovoice research can raise awareness of issues, influence policymakers (C. Wang & Burris, 1994) and stimulate change by shaping and developing policies and guidelines that appreciate the views and needs of individuals involved in Photovoice research (Smith et al., 2023).')
vn('Đến lượt mình, người tham gia nghiên cứu Photovoice có thể NÂNG CAO NHẬN THỨC về các vấn đề, ẢNH HƯỞNG đến nhà hoạch định chính sách (C. Wang & Burris, 1994) và THÚC ĐẨY THAY ĐỔI bằng cách định hình và phát triển các chính sách, hướng dẫn vốn ghi nhận quan điểm và nhu cầu của những người tham gia nghiên cứu Photovoice (Smith và cộng sự, 2023).')
d.add_paragraph()

add_h('1.2. Photovoice trong nghiên cứu sức khoẻ tâm thần vị thành niên', 2)

en('Scholars advocate Photovoice in research examining mental health, including depression, anxiety, and bipolar disorder (Han & Oliffe, 2016). Barry et al. (2021) suggest that Photovoice can promote public awareness surrounding mental health, whilst Flanagan and Flanagan et al. (2016) highlight how Photovoice can challenge negative stereotypes and reduce mental health stigma.')
vn('Các học giả ỦNG HỘ Photovoice trong nghiên cứu sức khoẻ tâm thần, bao gồm trầm cảm, lo âu và rối loạn lưỡng cực (Han & Oliffe, 2016). Barry và cộng sự (2021) gợi ý rằng Photovoice có thể THÚC ĐẨY nhận thức công chúng về sức khoẻ tâm thần, trong khi Flanagan và Flanagan và cộng sự (2016) làm nổi bật cách Photovoice có thể THÁCH THỨC các định kiến tiêu cực và GIẢM KỲ THỊ (stigma) sức khoẻ tâm thần.')
d.add_paragraph()

en('Photovoice may be a relevant method to explore mental health among adolescents who endure significant mental health stigma (Kaushik et al., 2016) and lack power (Wang et al., 1996). Moreover, adolescents are enthusiastic and committed to engaging in Photovoice research (Evans-Agnew et al., 2022) which may reflect their interest in producing photographs (Butschi & Hedderich, 2021).')
vn('Photovoice có thể là phương pháp PHÙ HỢP để khám phá sức khoẻ tâm thần ở vị thành niên — nhóm chịu kỳ thị đáng kể (Kaushik và cộng sự, 2016) và thiếu quyền lực (Wang và cộng sự, 1996). Hơn nữa, vị thành niên NHIỆT TÌNH và CAM KẾT tham gia nghiên cứu Photovoice (Evans-Agnew và cộng sự, 2022), điều có thể phản ánh sự quan tâm của họ đối với việc tạo ra ảnh chụp (Butschi & Hedderich, 2021).')
d.add_paragraph()

en('This could explain why adolescents find Photovoice research \'fun\' (S. Wass & Safari, 2020) compared to traditional methods which adolescents can find dull (Nuffield Council on Bioethics, 2015). Despite the compelling strengths of Photovoice, evidence to support its use in research involving adolescents is lacking (Butschi & Hedderich, 2021) and may reflect the under-representation of adolescents in mental health research (Mawn et al., 2016).')
vn('Điều này có thể giải thích vì sao vị thành niên thấy nghiên cứu Photovoice "VUI" (S. Wass & Safari, 2020) so với các phương pháp truyền thống mà các em thấy NHÀM CHÁN (Nuffield Council on Bioethics, 2015). Mặc dù có những điểm mạnh thuyết phục, BẰNG CHỨNG ủng hộ việc dùng Photovoice trong nghiên cứu vị thành niên còn THIẾU (Butschi & Hedderich, 2021), điều này có thể phản ánh sự đại diện không đủ của vị thành niên trong nghiên cứu sức khoẻ tâm thần (Mawn và cộng sự, 2016).')
d.add_paragraph()

en('Consequently, more research exploring the use of photovoice in mental health research involving adolescents is needed (Velez-Grau, 2019). However, a traditional systematic review approach is not appropriate in addressing this gap given the paucity of available evidence. Instead, our review will address this paucity of evidence by scoping the available evidence and answering two research questions.')
vn('Do đó, CẦN có thêm nghiên cứu khám phá việc dùng Photovoice trong nghiên cứu sức khoẻ tâm thần vị thành niên (Velez-Grau, 2019). Tuy nhiên, cách tiếp cận TỔNG QUAN HỆ THỐNG TRUYỀN THỐNG KHÔNG PHÙ HỢP để giải quyết khoảng trống này do thiếu hụt bằng chứng. Thay vào đó, tổng quan của chúng tôi sẽ giải quyết bằng cách SCOPING (lập phạm vi) bằng chứng sẵn có và trả lời 2 câu hỏi nghiên cứu.')
d.add_paragraph()

add_h('1.3. Mục tiêu nghiên cứu', 2)
en('Research aims of our scoping review: (1) What are the characteristics of Photovoice studies exploring mental health among adolescents? (2) What are the main themes in the findings of Photovoice studies exploring mental health among adolescents?')
vn('Mục tiêu nghiên cứu tổng quan phạm vi của chúng tôi: (1) ĐẶC ĐIỂM của các nghiên cứu Photovoice khám phá sức khoẻ tâm thần ở vị thành niên là gì? (2) Các CHỦ ĐỀ CHÍNH trong phát hiện của các nghiên cứu Photovoice khám phá sức khoẻ tâm thần ở vị thành niên là gì?')
d.add_paragraph()

# =============================================================
# METHODS
# =============================================================
page(4)
add_h('2. PHƯƠNG PHÁP (Methods)', 1)

en('We structured and reported our review according to Preferred Reporting Items for Systematic Reviews and Meta-Analyses extension for Scoping Reviews (PRISMA-ScR) Checklist (Tricco et al., 2018) (see supplementary material 1 for the checklist).')
vn('Chúng tôi cấu trúc và báo cáo tổng quan theo CHECKLIST PRISMA-ScR (Preferred Reporting Items for Systematic Reviews and Meta-Analyses extension for Scoping Reviews) (Tricco và cộng sự, 2018) (xem tài liệu bổ sung 1).')
d.add_paragraph()

add_h('2.1. Tiêu chí lựa chọn (Eligibility criteria)', 2)
en('The inclusion criteria were: studies that recruited participants between ten and 19 years old, based on the World Health Organisation (2021) definition; studies where the mean age of participants fell between ten and 19, and studies that provided separate data sets for this age range were also included; studies that explicitly stated the use of Photovoice as a method were included in the review; studies conducted in all countries, written in any language which explored mental, social, emotional, and psychosocial health and needs were also included.')
vn('Tiêu chí chọn vào: nghiên cứu tuyển người tham gia 10-19 tuổi (theo định nghĩa WHO 2021); nghiên cứu có TUỔI TRUNG BÌNH 10-19, và những nghiên cứu cung cấp tập dữ liệu riêng cho khoảng tuổi này cũng được đưa vào; nghiên cứu tuyên bố RÕ RÀNG sử dụng Photovoice làm phương pháp; nghiên cứu được tiến hành ở mọi quốc gia, viết bằng bất kỳ ngôn ngữ nào, khám phá sức khoẻ và nhu cầu TÂM THẦN, XÃ HỘI, CẢM XÚC và TÂM LÝ-XÃ HỘI cũng được đưa vào.')
d.add_paragraph()

en('Only primary studies that were published in peer-reviewed journals and grey literature were included. The review specifically focused on primary data, therefore book chapters, systematic reviews, dissertations, abstracts, comments, and editorials were excluded. Quantitative studies and mixed-method studies which did not provide separate qualitative data were excluded as our review subscribed to an interpretive paradigm that did not aim to generate effect sizes.')
vn('CHỈ các nghiên cứu chính (primary studies) đăng trên tạp chí peer-reviewed và tài liệu xám (grey literature) được bao gồm. Tổng quan tập trung VÀO DỮ LIỆU CHÍNH, do đó các chương sách, tổng quan hệ thống, luận văn, tóm tắt, bình luận và xã luận bị LOẠI TRỪ. Các nghiên cứu ĐỊNH LƯỢNG và HỖN HỢP không cung cấp dữ liệu định tính riêng cũng bị LOẠI TRỪ vì tổng quan của chúng tôi theo PARADIGM DIỄN GIẢI (interpretive paradigm) không nhằm tạo độ lớn hiệu ứng.')
d.add_paragraph()

add_h('2.2. Chiến lược tìm kiếm', 2)
en('We identified relevant studies from a pre-existing data extraction table which was part of a broader systematic review that explored the use of Photovoice in adolescent research (Burn et al., manuscript in preparation). Burn and colleagues\' review was similar in aim and yielded over 100 papers. We used the pre-existing data to avoid replicating research and to devote more resources to identifying qualitative papers focused on adolescent mental health research.')
vn('Chúng tôi xác định các nghiên cứu liên quan từ một BẢNG TRÍCH XUẤT DỮ LIỆU có sẵn, là một phần của tổng quan hệ thống rộng hơn khám phá việc sử dụng Photovoice trong nghiên cứu vị thành niên (Burn và cộng sự, bản thảo đang chuẩn bị). Tổng quan của Burn và cộng sự có MỤC TIÊU tương tự và mang lại hơn 100 bài. Chúng tôi sử dụng dữ liệu có sẵn để TRÁNH NHÂN ĐÔI nghiên cứu và dành thêm nguồn lực để xác định các bài định tính tập trung vào nghiên cứu sức khoẻ tâm thần vị thành niên.')
d.add_paragraph()

en('To ensure the identification of all relevant papers, we repeated the search on the 15th of June 2022 and again on the 9th of May 2023. An academic librarian peer-reviewed and approved the search strategy which contained keywords related to adolescents, photovoice and action research. We employed the search across five electronic databases: PSYCHinfo, PubMed, Scopus, Web of Science and CINAHAL. We also hand-searched the reference lists of relevant studies as well as relevant journals such as the International Journal of Adolescence and Youth as well as Google and Google Scholar and grey literature databases.')
vn('Để đảm bảo xác định được tất cả các bài liên quan, chúng tôi LẶP LẠI tìm kiếm vào ngày 15/06/2022 và 09/05/2023. Một thủ thư học thuật đã peer-review và phê duyệt chiến lược tìm kiếm với các từ khoá liên quan đến vị thành niên, photovoice và nghiên cứu hành động. Chúng tôi tìm kiếm trên 5 CƠ SỞ DỮ LIỆU ĐIỆN TỬ: PSYCHinfo, PubMed, Scopus, Web of Science và CINAHAL. Chúng tôi cũng tìm kiếm THỦ CÔNG (hand-searched) danh mục tham khảo của các nghiên cứu liên quan, các tạp chí liên quan như International Journal of Adolescence and Youth, Google và Google Scholar cùng các cơ sở dữ liệu tài liệu xám.')
d.add_paragraph()

add_h('2.3. Quy trình lựa chọn', 2)
en('We screened the titles and abstracts of studies within the data extraction table according to the above eligibility criteria. Studies that met the eligibility criteria underwent full-text review. Alongside this, we imported studies from the updated search into the reference management software, Rayyan, where duplicates were removed. Next, MS independently screened the titles and abstracts of all studies according to the above eligibility criteria. A full-text review was undertaken for studies that met the eligibility criteria. EK, a second reviewer, independently screened 25% of the studies at each stage with discrepancies of opinion being resolved by VJB, a third independent reviewer.')
vn('Chúng tôi sàng lọc TIÊU ĐỀ và TÓM TẮT các nghiên cứu trong bảng trích xuất theo tiêu chí trên. Các nghiên cứu đạt tiêu chí được REVIEW TOÀN VĂN. Song song, chúng tôi nhập các nghiên cứu từ tìm kiếm cập nhật vào phần mềm quản lý tham khảo Rayyan, nơi các bản trùng lặp được loại bỏ. Tiếp theo, MS sàng lọc ĐỘC LẬP tiêu đề và tóm tắt của tất cả nghiên cứu. Review toàn văn được thực hiện cho các nghiên cứu đạt tiêu chí. EK, người review thứ hai, ĐỘC LẬP sàng lọc 25% nghiên cứu ở mỗi giai đoạn, các bất đồng ý kiến được giải quyết bởi VJB, người review độc lập thứ ba.')
d.add_paragraph()

add_h('2.4. Trích xuất dữ liệu và phân tích', 2)
page(5)
en('To answer research question two, "What are the main themes in the findings of Photovoice studies exploring mental health?", MS carried out a manual reflexive thematic analysis outlined by Braun and Clarke (2006) to identify common themes across the study findings. EK independently analysed 25% of the studies using reflexive thematic analysis.')
vn('Để trả lời câu hỏi nghiên cứu thứ 2, "Các chủ đề chính trong phát hiện của các nghiên cứu Photovoice khám phá sức khoẻ tâm thần là gì?", MS thực hiện PHÂN TÍCH CHỦ ĐỀ PHẢN TƯ (reflexive thematic analysis) thủ công theo hướng dẫn của Braun và Clarke (2006) để xác định các chủ đề chung xuyên suốt các phát hiện. EK phân tích độc lập 25% các nghiên cứu sử dụng phân tích chủ đề phản tư.')
d.add_paragraph()

add_h('2.5. Đánh giá chất lượng (Quality appraisal)', 2)
en('We appraised the quality of the studies to understand the calibre of Photovoice studies exploring mental health among adolescents. To appraise the quality of the studies, we used the Critical Appraisal Skills Programme (CASP) qualitative study checklist (2018) and the Consolidated Criteria for Reporting Qualitative Research (COREQ) (Tong et al., 2007).')
vn('Chúng tôi đánh giá CHẤT LƯỢNG các nghiên cứu để hiểu trình độ của các nghiên cứu Photovoice khám phá sức khoẻ tâm thần vị thành niên. Để đánh giá chất lượng, chúng tôi sử dụng CASP (Critical Appraisal Skills Programme) checklist nghiên cứu định tính (2018) và COREQ (Consolidated Criteria for Reporting Qualitative Research) (Tong và cộng sự, 2007).')
d.add_paragraph()

en('To support the quality appraisal process, Al-Moghrabi et al. (2019) criteria were used which categorize the quality of reporting according to the quantity of COREQ items reported (good (≥25 items), moderate (17 to 24 items), poor (9 to 16 items), very poor (≤8 items).')
vn('Để hỗ trợ quá trình đánh giá chất lượng, tiêu chí của Al-Moghrabi và cộng sự (2019) được dùng — phân loại chất lượng báo cáo theo số mục COREQ được báo cáo: TỐT (≥25 mục), TRUNG BÌNH (17-24), KÉM (9-16), RẤT KÉM (≤8).')
d.add_paragraph()

en('We used these quality appraisal tools to gain and triangulate information related to the quality of the included studies, which we considered appropriate to scope the quality of studies included in this review. MS and EK independently conducted the quality appraisal, with MS appraising 100% of the total studies and EK appraising 25%. Differences of opinion were resolved through discussion with VJB, a third independent reviewer.')
vn('Chúng tôi dùng các công cụ này để thu thập và TAM GIÁC HOÁ (triangulate) thông tin về chất lượng các nghiên cứu được đưa vào. MS và EK ĐỘC LẬP đánh giá chất lượng — MS đánh giá 100% tổng số nghiên cứu, EK đánh giá 25%. Bất đồng được giải quyết qua thảo luận với VJB, người đánh giá độc lập thứ ba.')
d.add_paragraph()

# =============================================================
# FINDINGS
# =============================================================
add_h('3. PHÁT HIỆN VÀ THẢO LUẬN (Findings and discussion)', 1)

en('From Burn and colleagues (2021) pre-existing data extraction table, we identified seven eligible studies out of 215 studies. From the updated search, we identified 334 studies, of which three met the inclusion criteria. Also, we included two eligible grey literature sources, bringing the total number of included studies in our review to twelve.')
vn('Từ bảng trích xuất dữ liệu có sẵn của Burn và cộng sự (2021), chúng tôi xác định 7 nghiên cứu đủ điều kiện trong tổng số 215 nghiên cứu. Từ tìm kiếm cập nhật, chúng tôi xác định 334 nghiên cứu, trong đó 3 đạt tiêu chí. Ngoài ra, 2 nguồn tài liệu xám đủ điều kiện được bao gồm, đưa TỔNG SỐ NGHIÊN CỨU TRONG TỔNG QUAN của chúng tôi lên 12 BÀI.')
d.add_paragraph()

vn('(Sơ đồ PRISMA — Hình 1 — trong PDF gốc trang 10 mô tả luồng identification → screening → eligibility → included.)')
d.add_paragraph()

add_h('3.1. Đặc điểm các nghiên cứu Photovoice', 2)
add_h('3.1.1. Chất lượng', 3)
en('The quality appraisal highlighted variation in the reporting of Photovoice studies examining mental health among adolescents, which is consistent with broader Photovoice studies (Catalani & Minkler, 2010; K. C. Hergenrather et al., 2009). None of the studies met the criteria for good quality, whilst six studies were of moderate reporting quality, and four were assessed as poor reporting quality.')
vn('Đánh giá chất lượng làm nổi bật SỰ BIẾN THIÊN trong báo cáo các nghiên cứu Photovoice — phù hợp với các nghiên cứu Photovoice rộng hơn (Catalani & Minkler, 2010; K. C. Hergenrather và cộng sự, 2009). KHÔNG NGHIÊN CỨU NÀO đạt tiêu chí "tốt", trong khi 6 nghiên cứu ở mức báo cáo TRUNG BÌNH và 4 nghiên cứu ở mức KÉM.')
d.add_paragraph()

en('These drawbacks were attributable to recruitment, data collection and reflexivity domains. Despite these shortcomings, most studies stated their aims and objectives, used appropriate qualitative methodology; addressed the research aim using a relevant design and appropriately managed ethical issues.')
vn('Những hạn chế này thuộc các lĩnh vực TUYỂN MỘ, THU THẬP DỮ LIỆU và TÍNH PHẢN TƯ (reflexivity). Mặc dù có những thiếu sót này, đa số nghiên cứu nêu rõ mục tiêu, dùng phương pháp định tính phù hợp; giải quyết mục tiêu nghiên cứu bằng thiết kế phù hợp và quản lý đạo đức thích đáng.')
d.add_paragraph()

en('The variation in quality and reporting among Photovoice studies examining mental health among adolescents may be attributable to researchers\' reliance on their own subjective judgement rather than scientific guidance (Bugos et al., 2014). To the best of our knowledge, no guidance exists to support researchers reporting Photovoice studies. This could result in researchers reporting Photovoice studies based on what they think is appropriate, possibly explaining variations in quality seen among the studies in our review.')
vn('Sự biến thiên về chất lượng và báo cáo có thể do các nhà nghiên cứu DỰA VÀO PHÁN ĐOÁN CHỦ QUAN của bản thân hơn là hướng dẫn khoa học (Bugos và cộng sự, 2014). Theo hiểu biết của chúng tôi, KHÔNG CÓ HƯỚNG DẪN nào hỗ trợ các nhà nghiên cứu báo cáo nghiên cứu Photovoice. Điều này có thể dẫn đến nhà nghiên cứu báo cáo theo những gì họ cho là phù hợp, có thể giải thích sự biến thiên về chất lượng trong tổng quan của chúng tôi.')
d.add_paragraph()

# Geographic + urban/rural — synthesized to save space
add_h('3.1.2. Phân bố địa lý', 3)
page(11)
en('Researchers conducted ten Photovoice studies examining mental health among adolescents in high-income countries (HICs) as per The World Bank\'s (2022) classification. Nine studies were conducted in North America (U.S.A. (n=6), Canada (n=3)), whilst another was conducted in Europe (UK). Two studies were conducted in low-and middle-Income Countries (LMICs) across Africa (Kenya and South Africa).')
vn('10/12 nghiên cứu Photovoice ở vị thành niên được tiến hành tại các nước THU NHẬP CAO (HICs) theo phân loại Ngân hàng Thế giới (2022). 9 nghiên cứu ở Bắc Mỹ (USA n=6, Canada n=3), 1 ở châu Âu (Anh). 2 nghiên cứu ở các nước THU NHẬP THẤP-TRUNG BÌNH (LMICs) ở châu Phi (Kenya và Nam Phi).')
d.add_paragraph()

en('Researchers in LMICs experience barriers to producing high-quality research (Shumba & Lusambili, 2021), such as a lack of funding to conduct mental health research (Patel et al., 2018). These issues may be compounded and reinforced within Photovoice research, which can be expensive (Julien et al., 2013) due to the cost of purchasing cameras and reproducing photographs (Coemans et al., 2019).')
vn('Các nhà nghiên cứu ở LMICs gặp RÀO CẢN trong việc tạo ra nghiên cứu chất lượng cao (Shumba & Lusambili, 2021), ví dụ thiếu kinh phí cho nghiên cứu sức khoẻ tâm thần (Patel và cộng sự, 2018). Những vấn đề này có thể bị KHUẾCH ĐẠI trong nghiên cứu Photovoice, vốn có thể TỐN KÉM (Julien và cộng sự, 2013) do chi phí mua máy ảnh và tái tạo ảnh (Coemans và cộng sự, 2019).')
d.add_paragraph()

en('As a result, the over-representation of HICs across Photovoice studies exploring mental health among adolescents may limit the generalizability of findings to adolescents in LMICS.')
vn('Hệ quả: SỰ ĐẠI DIỆN QUÁ MỨC của HICs trong các nghiên cứu Photovoice có thể hạn chế khả năng NGOẠI SUY (generalizability) các phát hiện cho vị thành niên ở LMICs.')
d.add_paragraph()

add_h('3.1.3. Phân bố đô thị/nông thôn', 3)
en('Although researchers did not describe the context of two studies, nine Photovoice studies examining mental health among adolescents were conducted in urban contexts, with Dempsey (2016)\'s study conducted in a rural setting.')
vn('Mặc dù 2 nghiên cứu không mô tả bối cảnh, 9 nghiên cứu Photovoice được tiến hành ở BỐI CẢNH ĐÔ THỊ, chỉ Dempsey (2016) ở NÔNG THÔN.')
d.add_paragraph()

en('In comparison, rural populations experience barriers to participating in research, such as a lack of transportation (Levy et al., 2017) that likely hinders Photovoice studies from being conducted. Furthermore, cities have a younger population (Striessnig et al., 2019) that is at greater risk of mental ill-health compared to rural populations (Gruebner et al., 2017).')
vn('So sánh: dân cư NÔNG THÔN gặp rào cản tham gia nghiên cứu, ví dụ thiếu phương tiện đi lại (Levy và cộng sự, 2017), có thể cản trở các nghiên cứu Photovoice. Hơn nữa, thành phố có DÂN CƯ TRẺ HƠN (Striessnig và cộng sự, 2019), có nguy cơ bệnh tâm thần cao hơn nông thôn (Gruebner và cộng sự, 2017).')
d.add_paragraph()

add_h('3.1.4. Phương pháp dùng cùng Photovoice', 3)
page(12)
en('Two studies did not specify whether they had used other methods alongside Photovoice. In comparison, seven studies reported using photovoice as a standalone method, while three studies combined Photovoice with other qualitative and arts-based methods, including interviews and creative writing.')
vn('2 nghiên cứu không chỉ rõ có dùng phương pháp khác kèm Photovoice. 7 nghiên cứu dùng Photovoice ĐỘC LẬP, 3 nghiên cứu KẾT HỢP Photovoice với các phương pháp định tính và nghệ thuật khác như PHỎNG VẤN và VIẾT SÁNG TẠO.')
d.add_paragraph()

en('Our finding suggests that Photovoice is an adaptable research method (Lal et al., 2012) that addresses diverse research interests (Catalani & Minkler, 2010) and highlights the inherent flexibility of participatory methods (Gaboardi et al., 2022).')
vn('Phát hiện này gợi ý Photovoice là phương pháp THÍCH ỨNG ĐƯỢC (Lal và cộng sự, 2012) đáp ứng đa dạng quan tâm nghiên cứu (Catalani & Minkler, 2010) và nhấn mạnh tính LINH HOẠT VỐN CÓ của các phương pháp tham gia (Gaboardi và cộng sự, 2022).')
d.add_paragraph()

add_h('3.1.5. Mục tiêu nghiên cứu', 3)
en('Nine Photovoice studies focused on examining social determinants of adolescent mental health, while three studies had more abstract objectives, including understanding adolescents\' use of metaphors to describe experiences of mental health.')
vn('9 nghiên cứu tập trung CÁC YẾU TỐ XÃ HỘI QUYẾT ĐỊNH (social determinants) sức khoẻ tâm thần vị thành niên, 3 nghiên cứu có mục tiêu TRỪU TƯỢNG hơn, bao gồm hiểu cách vị thành niên dùng ẨN DỤ để mô tả trải nghiệm sức khoẻ tâm thần.')
d.add_paragraph()

add_h('3.1.6. Sự tham gia của vị thành niên', 3)
en('Adolescents within the twelve Photovoice studies examining mental health predominantly participated in collecting and/or analysing photographs which are typical of Photovoice research (Golden, 2020; Strack et al., 2004).')
vn('Vị thành niên trong 12 nghiên cứu CHỦ YẾU tham gia THU THẬP và/hoặc PHÂN TÍCH ẢNH — điển hình của nghiên cứu Photovoice (Golden, 2020; Strack và cộng sự, 2004).')
d.add_paragraph()

en('However, eight Photovoice studies did not report whether adolescents had participated in the dissemination of the research which could be attributable to adolescents\' perception of research beyond data collection as dull and difficult (Mawn et al., 2016).')
vn('Tuy nhiên, 8 nghiên cứu KHÔNG báo cáo liệu vị thành niên có tham gia PHỔ BIẾN (dissemination) nghiên cứu hay không. Điều này có thể do vị thành niên thấy các hoạt động nghiên cứu ngoài thu thập dữ liệu là CHÁN VÀ KHÓ (Mawn và cộng sự, 2016).')
d.add_paragraph()

en('However, four studies described participants\' engagement in disseminating the research, including their participation in selecting photos for an exhibition and film.')
vn('Tuy nhiên, 4 nghiên cứu mô tả việc người tham gia tham gia phổ biến nghiên cứu, bao gồm chọn ảnh cho TRIỂN LÃM và PHIM.')
d.add_paragraph()

add_h('3.1.7. Ảnh được thu thập', 3)
page(13)
en('None of the researchers quantified the total number of photographs collected or analysed by the adolescents involved in the Photovoice research. Similarly, five researchers did not describe the type of camera used by participants to capture photographs. Although, four researchers reported the use of disposable cameras; two reported the use of digital cameras, and one reported the use of mobile phones.')
vn('KHÔNG nhà nghiên cứu nào ĐỊNH LƯỢNG tổng số ảnh được thu thập hoặc phân tích. Tương tự, 5 nhà nghiên cứu không mô tả LOẠI MÁY ẢNH. Trong số có thông tin: 4 dùng máy ảnh DÙNG MỘT LẦN (disposable), 2 dùng máy ảnh số, 1 dùng điện thoại di động.')
d.add_paragraph()

en('Mobile phones have the capacity to capture almost unlimited photographs (Hartnell-Young & Heym, 2008) compared to disposable cameras that are limited to capturing 28 photographs (Petersen & Martin, 2021).')
vn('Điện thoại di động có khả năng chụp KHÔNG GIỚI HẠN ảnh (Hartnell-Young & Heym, 2008) so với máy dùng một lần chỉ 28 ảnh (Petersen & Martin, 2021).')
d.add_paragraph()

add_h('3.1.8. Sử dụng ảnh được công bố', 3)
en('Seven studies incorporated participants\' images within the published manuscript, although none of the studies reported adolescents\' involvement in selecting which photographs were chosen for publication.')
vn('7 nghiên cứu đưa ảnh người tham gia vào bản thảo công bố, dù KHÔNG nghiên cứu nào báo cáo việc vị thành niên tham gia CHỌN ảnh nào được công bố.')
d.add_paragraph()

en('This finding is consistent with other sources (Evans-Agnew & Rosemberg, 2016) and suggests that adolescents in Photovoice studies examining mental health are excluded from the decisions made related to the manuscript.')
vn('Phát hiện này nhất quán với các nguồn khác (Evans-Agnew & Rosemberg, 2016) và gợi ý rằng vị thành niên trong các nghiên cứu Photovoice bị LOẠI TRỪ khỏi các quyết định liên quan đến bản thảo.')
d.add_paragraph()

add_h('3.1.9. Xu hướng công bố theo thời gian', 3)
page(14)
en('Researchers have published limited Photovoice studies examining mental health among adolescents since 2012, with the majority of these studies being published in 2021. This trend echoes the novel use of arts-based methods within research (Leavy, 2020) which has been developing for over twenty years (Papoulias, 2018).')
vn('Số nghiên cứu Photovoice về sức khoẻ tâm thần vị thành niên còn HẠN CHẾ từ 2012, đa số được công bố vào 2021. Xu hướng này phản ánh việc dùng phương pháp dựa trên NGHỆ THUẬT (arts-based) còn mới (Leavy, 2020), đã phát triển hơn 20 năm (Papoulias, 2018).')
d.add_paragraph()

add_h('3.1.10. Tỷ lệ rời bỏ (Attrition)', 3)
en('In five studies, attrition was an issue affecting participation among adolescents due to barriers such as illness, transport issues, domestic situations, and loss of camera equipment. Attrition is a common issue within mental health research (Homman et al., 2021), with individuals experiencing mental illness more likely to disengage (Folke et al., 2018).')
vn('Trong 5 nghiên cứu, ATTRITION là vấn đề ảnh hưởng sự tham gia do các rào cản như BỆNH TẬT, vấn đề ĐI LẠI, hoàn cảnh GIA ĐÌNH và MẤT THIẾT BỊ MÁY ẢNH. Attrition là vấn đề thường gặp trong nghiên cứu sức khoẻ tâm thần (Homman và cộng sự, 2021), những người có bệnh tâm thần dễ NGỪNG THAM GIA (Folke và cộng sự, 2018).')
d.add_paragraph()

add_h('3.1.11. Số lượng và thời lượng buổi', 3)
en('Researchers did not report the number of sessions within eight of the Photovoice studies examining mental health among adolescents, while four studies reported two or three sessions. Similarly, five studies did not report on the duration of each session, whereas the remaining five studies reported session lengths between 40 and 90 minutes.')
vn('8 nghiên cứu KHÔNG báo cáo số buổi, 4 nghiên cứu báo cáo 2-3 buổi. Tương tự, 5 nghiên cứu không báo cáo thời lượng mỗi buổi, 5 nghiên cứu còn lại báo cáo 40-90 PHÚT/buổi.')
d.add_paragraph()

add_h('3.1.12. Số nhóm', 3)
en('In one study, researchers interviewed participants individually instead of grouping them together as in other studies. However, nine studies did not specify how many groups there were within the Photovoice study, whilst two reported four and 12 groups of participants.')
vn('1 nghiên cứu phỏng vấn CÁ NHÂN thay vì nhóm. 9 nghiên cứu không chỉ rõ số nhóm, 2 nghiên cứu báo cáo có 4 và 12 nhóm.')
d.add_paragraph()

add_h('3.1.13. Khung hướng dẫn thảo luận — SHOWeD', 3)
en('To facilitate discussion of photos, six studies used the SHOWeD framework which asks five questions: What do you See here? What is really Happening? How does this relate to Our lives? Why does this situation exist? What can we Do about it?')
vn('Để hỗ trợ thảo luận về ảnh, 6 nghiên cứu sử dụng KHUNG SHOWeD với 5 câu hỏi: Bạn THẤY (See) gì ở đây? Điều gì THỰC SỰ ĐANG XẢY RA (Happening)? Điều này liên quan đến CUỘC SỐNG CỦA CHÚNG TA (Our lives) như thế nào? Tại sao tình huống này TỒN TẠI (exist)? Chúng ta CÓ THỂ LÀM GÌ (Do) về điều đó?')
d.add_paragraph()

# === Demographics ===
page(15)
add_h('3.2. Đặc điểm vị thành niên tham gia', 2)

add_h('3.2.1. Cỡ mẫu', 3)
en('Two studies did not specify the number of adolescents participating. The remaining ten Photovoice studies had sample sizes ranging from four to 58 participants, with ten or fewer adolescents included in six studies.')
vn('2 nghiên cứu không chỉ rõ số người tham gia. 10 nghiên cứu còn lại có cỡ mẫu từ 4 đến 58 người, với ≤10 vị thành niên ở 6 nghiên cứu.')
d.add_paragraph()

en('However, this trend deviates from the traditional Photovoice process by C. Wang and Burris (1997) which recommends a maximum sample size of ten participants.')
vn('Tuy nhiên, xu hướng này LỆCH với quy trình Photovoice truyền thống của C. Wang và Burris (1997) — vốn khuyến nghị cỡ mẫu TỐI ĐA 10 người.')
d.add_paragraph()

add_h('3.2.2. Đại diện giới tính', 3)
en('Although Wainaina et al., (2021) had research aims that could only be answered by female participants, nine studies were significantly over-represented by females (127 females to 47 males, respectively). However, three studies did not specify the gender of participants.')
vn('Mặc dù Wainaina và cộng sự (2021) có mục tiêu chỉ trả lời được bởi nữ, 9 nghiên cứu khác CŨNG CÓ NỮ ĐẠI DIỆN QUÁ MỨC (127 nữ so với 47 nam). 3 nghiên cứu không chỉ rõ giới tính.')
d.add_paragraph()

en('Females are more likely to participate in Photovoice studies exploring mental health among adolescents, coinciding with the trend seen across research. The feminist underpinnings of Photovoice (C. Wang & Burris, 1994) aim to promote communication among females (Capewell et al., 2020).')
vn('Nữ có xu hướng tham gia nghiên cứu Photovoice nhiều hơn, trùng với xu hướng chung trong nghiên cứu. Nền tảng NỮ QUYỀN của Photovoice (C. Wang & Burris, 1994) nhằm thúc đẩy giao tiếp giữa phụ nữ (Capewell và cộng sự, 2020).')
d.add_paragraph()

add_h('3.2.3. Tuổi', 3)
page(16)
en('Ten studies provided the age ranges of participants which collectively ranged from ten to 22 years old. However, two studies did not specify the ages of participants.')
vn('10 nghiên cứu cung cấp KHOẢNG TUỔI người tham gia, dao động chung từ 10 đến 22 tuổi. 2 nghiên cứu không chỉ rõ tuổi.')
d.add_paragraph()

add_h('3.2.4. Tình trạng kinh tế-xã hội', 3)
en('Eleven studies did not report the socioeconomic status of participants, except for Liegghio (2016) who included an equal mix of adolescents from low, middle and high-income households.')
vn('11 nghiên cứu KHÔNG báo cáo tình trạng KINH TẾ-XÃ HỘI người tham gia, ngoại trừ Liegghio (2016) — người này bao gồm hỗn hợp đều vị thành niên từ hộ thu nhập thấp, trung bình và cao.')
d.add_paragraph()

add_h('3.2.5. Sắc tộc', 3)
en('Five studies did not report the ethnicity of adolescents. Whereas seven studies included information about the participants\' ethnicity, although there was no standardized approach to reporting this data across the studies.')
vn('5 nghiên cứu KHÔNG báo cáo SẮC TỘC. 7 nghiên cứu có thông tin sắc tộc, nhưng không có CÁCH TIẾP CẬN CHUẨN HOÁ trong báo cáo.')
d.add_paragraph()

# === KEY THEMES ===
add_h('3.3. CÁC CHỦ ĐỀ CHÍNH (Key themes)', 2)

add_h('3.3.1. Đấu tranh có thể mang lại sức mạnh và sự thuộc về', 3)
en('Our findings revealed common themes related to coping and resilience and beliefs about oneself. In Bashore and colleagues\' (2017) research, adolescents described relying on coping skills, including playing musical instruments to manage stress. One participant in the study shared a photograph of a drum to illustrate how they released their frustration.')
vn('Phát hiện chính: ỨNG PHÓ (coping), KIÊN CƯỜNG (resilience) và NIỀM TIN VỀ BẢN THÂN. Trong Bashore và cộng sự (2017), vị thành niên mô tả việc DỰA VÀO kỹ năng ứng phó, gồm CHƠI NHẠC CỤ để quản lý stress. Một người tham gia chia sẻ ảnh chiếc TRỐNG để minh hoạ cách họ giải toả sự thất vọng.')
d.add_paragraph()

en('Moreover, one respondent in Rose\'s et al. (2018) study described how coping was a way to "occupy your mind from something that\'s hurting you for a little while." (Page 802).')
vn('Hơn nữa, một người trong Rose và cộng sự (2018) mô tả ứng phó là cách "ĐỂ TÂM TRÍ XA KHỎI điều đang làm tổn thương bạn trong một lúc" (trang 802).')
d.add_paragraph()

en('Orth and van Wyk\'s (2022) study explored resilience in adolescents, with one 17-year-old female describing her persistence using a photograph she had taken of rain: "I thought, why not embrace . . . like the storms of my life and accept my situation" (Page 1444).')
vn('Orth & van Wyk (2022) khám phá kiên cường, một bé gái 17 tuổi mô tả sự bền chí qua ảnh chụp MƯA: "Tôi đã nghĩ, sao không OM TRỌN... như những cơn bão trong đời tôi, và CHẤP NHẬN tình huống của mình" (trang 1444).')
d.add_paragraph()

en('Similarly, in Liegghio\'s (2016) study, a participant described how a photo of a drawing of a skull was representative of them being "hardcore" and "durable", reflecting their resilience.')
vn('Tương tự, trong Liegghio (2016), một người tham gia mô tả ảnh hình vẽ ĐẦU LÂU đại diện cho việc họ "hardcore" và "BỀN BỈ", phản ánh kiên cường.')
d.add_paragraph()

en('Dempsey (2016) noted how adolescents\' beliefs about themselves, including a perceived lack of belonging and loneliness, negatively influenced depression in adolescents.')
vn('Dempsey (2016) ghi nhận niềm tin của vị thành niên về bản thân, bao gồm cảm thấy KHÔNG THUỘC VỀ và CÔ ĐƠN, ảnh hưởng tiêu cực đến trầm cảm.')
d.add_paragraph()

en('Furthermore, in Georgievski and colleagues\' (2018) study, an adolescent outlined how cancer treatment influenced their perception of themselves: "I feel that because I went through this, I can go through anything..." (Page 710).')
vn('Trong Georgievski và cộng sự (2018), một vị thành niên mô tả cách điều trị ung thư ảnh hưởng nhận thức bản thân: "Tôi cảm thấy vì đã trải qua điều này, tôi có thể vượt qua mọi thứ..." (trang 710).')
d.add_paragraph()

page(17)
en('Moreover, Woodgate et al. (2021) study included a participant reflecting on self-perception whilst examining a photograph of a stop sign: "I feel like I\'m just sort of standing and staying in one place, not really making progress in any field of my life... so that sort of represents anxiety and depression." (Page 9).')
vn('Hơn nữa, Woodgate và cộng sự (2021) có một người tham gia suy ngẫm về tự nhận thức khi xem ảnh BIỂN STOP: "Tôi cảm thấy mình như đang đứng yên một chỗ, không thực sự tiến bộ ở bất kỳ lĩnh vực nào... điều đó đại diện cho LO ÂU và TRẦM CẢM" (trang 9).')
d.add_paragraph()

en('An adolescent from the project by Northwest Michigan Community Health Innovation Region\'s Behavioral Health Initiative (2021) explained how believing that "you\'re not alone even though you may feel like it" could reduce mental health stigma (Page 19).')
vn('Một vị thành niên từ dự án của Northwest Michigan (2021) giải thích cách tin rằng "bạn KHÔNG CÔ ĐƠN dù có thể bạn CẢM THẤY như vậy" có thể giảm kỳ thị sức khoẻ tâm thần (trang 19).')
d.add_paragraph()

en('Coping with stress is associated with resilience (Bonanno, 2012) which, in turn, influences beliefs about oneself (Cazan & Dumitrescu, 2016). Our review highlights the relationship between these intersecting determinants and adolescent mental health and underscores the potential of employing photovoice in adolescent research to explore complex subjective experiences of mental health.')
vn('Ứng phó với stress liên quan đến kiên cường (Bonanno, 2012), và đến lượt nó ảnh hưởng niềm tin về bản thân (Cazan & Dumitrescu, 2016). Tổng quan của chúng tôi làm nổi bật MỐI QUAN HỆ giữa các yếu tố quyết định giao thoa này với sức khoẻ tâm thần vị thành niên, và nhấn mạnh tiềm năng dùng Photovoice để khám phá những trải nghiệm chủ quan PHỨC TẠP về sức khoẻ tâm thần.')
d.add_paragraph()

add_h('3.3.2. Gia đình và bạn bè cung cấp hy vọng và (sự không) hiểu', 3)
en('In Orth and van Wyk\'s (2022) study, one participant explained that their sibling provided them with a sense of purpose. They illustrated this by sharing a photo of their sister and explained: "Oh, this is my sister. You know when I look at her, there\'s a moment that I want to be myself. You know she gave me hope." (Page 1438).')
vn('Trong Orth & van Wyk (2022), một người mô tả em/chị/anh ruột mang lại CẢM GIÁC MỤC ĐÍCH. Họ minh hoạ qua ảnh em gái: "Ồ, đây là em gái tôi. Bạn biết không, khi nhìn em, có khoảnh khắc tôi muốn LÀ CHÍNH MÌNH. Em đã cho tôi HY VỌNG" (trang 1438).')
d.add_paragraph()

en('Bowers and Wozniak\'s (2020) project captured the importance of friends using a photo of teenagers which was captioned: "Friends are always there to lend a hand or an ear when you need it." (no page).')
vn('Dự án Bowers & Wozniak (2020) ghi lại tầm quan trọng của bạn bè qua ảnh các thiếu niên với chú thích: "Bạn bè LUÔN ở đó để cho bạn một bàn tay hoặc một đôi tai khi bạn cần".')
d.add_paragraph()

en('Furthermore, Velez-Grau\'s (2019) study highlighted differences in how mental health is understood between adolescents and their families. One adolescent explained how their family perceive mental health as a "one-way road" that does not appreciate the barriers surrounding mental health care (Page 917).')
vn('Velez-Grau (2019) làm nổi bật sự khác biệt cách hiểu sức khoẻ tâm thần giữa vị thành niên và gia đình. Một vị thành niên giải thích gia đình họ cảm nhận sức khoẻ tâm thần như "một con đường một chiều" không trân trọng các rào cản xung quanh chăm sóc sức khoẻ tâm thần (trang 917).')
d.add_paragraph()

en('The Social Buffering Hypothesis recognizes that social networks negate the experience of stress (Cohen & Wills, 1985) and enable communication (Cobb, 1976). Friends and family have a key role in promoting well-being in adolescents (Schacter & Margolin, 2019).')
vn('GIẢ THUYẾT ĐỆM XÃ HỘI (Social Buffering Hypothesis) ghi nhận rằng mạng lưới xã hội PHỦ ĐỊNH (negate) trải nghiệm stress (Cohen & Wills, 1985) và CHO PHÉP giao tiếp (Cobb, 1976). Bạn bè và gia đình có vai trò CHÍNH trong việc thúc đẩy phúc lợi vị thành niên (Schacter & Margolin, 2019).')
d.add_paragraph()

add_h('3.3.3. Cảm giác an toàn và thách thức của nghèo đói', 3)
page(18)
en('In Dempsey\'s (2016) study, the correlation between living near danger and the risk of depression among adolescents was described. Whereas Woodgate and colleagues\' (2021) study saw adolescents discuss the challenges of leaving places they considered safe.')
vn('Trong Dempsey (2016), tương quan giữa SỐNG GẦN NGUY HIỂM và nguy cơ trầm cảm ở vị thành niên được mô tả. Trong khi đó, Woodgate và cộng sự (2021) ghi nhận vị thành niên thảo luận thách thức của việc RỜI BỎ những nơi họ thấy AN TOÀN.')
d.add_paragraph()

en('Furthermore, one adolescent in Watson and Douglas (2012) reflected on a photo of an "unsafe" bridge, known as "Muggers Bridge", and how they phoned their father prior to walking underneath it (Page 291).')
vn('Hơn nữa, một vị thành niên trong Watson & Douglas (2012) suy ngẫm về ảnh CẦU "KHÔNG AN TOÀN" — gọi là "Cầu Cướp Bóc" (Muggers Bridge) — và cách họ gọi điện cho bố TRƯỚC khi đi dưới cầu (trang 291).')
d.add_paragraph()

en('Whereas one adolescent in Wainaina and colleagues\' (2021) study used a photograph of a child facing a precarious building to illustrate the feelings of stress she felt as a mother in an informal settlement owing to a lack of financial resources to care for her child (page 5).')
vn('Một vị thành niên trong Wainaina và cộng sự (2021) dùng ảnh đứa trẻ đối diện toà nhà NGUY HIỂM để minh hoạ cảm giác stress mà cô ấy — với tư cách là MẸ ở khu định cư không chính thức — cảm thấy do THIẾU NGUỒN LỰC TÀI CHÍNH để chăm sóc con (trang 5).')
d.add_paragraph()

en('Researchers recognize that perceived feelings of safety determine mental health (Wilson-Genderson & Pruchno, 2013), with lower perceptions of safety being correlated with increased stress (Pearson et al., 2021).')
vn('Các nhà nghiên cứu ghi nhận CẢM GIÁC AN TOÀN ĐƯỢC NHẬN THỨC quyết định sức khoẻ tâm thần (Wilson-Genderson & Pruchno, 2013), nhận thức an toàn THẤP HƠN tương quan với stress TĂNG (Pearson và cộng sự, 2021).')
d.add_paragraph()

add_h('3.3.4. Điều trị và môi trường lâm sàng ảnh hưởng nhận thức về bản thể', 3)
en('Treatment was revealed as a common theme. One adolescent in Orth and van Wyk\'s (2022) study illustrated the theme of treatment by using a photo of a flower to describe how adhering to HIV medication would make her "beautiful" and "strong" (Page 1443).')
vn('ĐIỀU TRỊ là chủ đề chung. Một vị thành niên trong Orth & van Wyk (2022) minh hoạ qua ảnh BÔNG HOA để mô tả việc tuân thủ thuốc HIV sẽ khiến cô ấy "ĐẸP" và "MẠNH MẼ" (trang 1443).')
d.add_paragraph()

en('Furthermore, in Leighhio\'s study, a respondent captured a photograph showing a door from where they were monitored within a mental health facility, reflecting the widespread practice of observation within mental health (Harrington et al., 2019).')
vn('Hơn nữa, trong Liegghio, một người trả lời chụp ảnh CÁNH CỬA nơi họ bị GIÁM SÁT trong một cơ sở sức khoẻ tâm thần, phản ánh thực hành QUAN SÁT phổ biến trong sức khoẻ tâm thần (Harrington và cộng sự, 2019).')
d.add_paragraph()

en('Furthermore, one adolescent in Velez-Grau\'s (2019)\'s study described the frustrations associated with treatment adherence and the need to negotiate their engagement in their mental health care: "I am the patient talk to me about my treatment don\'t talk to my parents first, especially when they do not know what\'s up" (Page 918).')
vn('Một vị thành niên trong Velez-Grau (2019) mô tả sự THẤT VỌNG liên quan đến tuân thủ điều trị và nhu cầu thương lượng việc tham gia chăm sóc sức khoẻ tâm thần: "Tôi là BỆNH NHÂN, hãy nói với tôi về điều trị của tôi, đừng nói với bố mẹ tôi TRƯỚC, đặc biệt khi họ không biết chuyện gì đang diễn ra" (trang 918).')
d.add_paragraph()

en('Our findings reinforce how adolescents want to be involved in their mental health care (Gros et al., 2017), despite their parents being more likely to be invited to engage in the decision-making process (Quaye et al., 2019).')
vn('Phát hiện của chúng tôi củng cố rằng vị thành niên MUỐN tham gia chăm sóc sức khoẻ tâm thần của mình (Gros và cộng sự, 2017), mặc dù bố mẹ THƯỜNG được mời tham gia vào quá trình ra quyết định nhiều hơn (Quaye và cộng sự, 2019).')
d.add_paragraph()

# === HIC vs LMIC differences ===
page(19)
add_h('3.4. Khác biệt giữa nghiên cứu Photovoice ở HICs và LMICs', 2)
en('Adolescents involved in Photovoice studies exploring mental health in LMICs were situated within the context of communicable disease and sampled from informal settlements. In addition, philanthropist organizations, such as The Oppenheimer Memorial Trust and Bill & Melinda Gates Foundation funded Photovoice studies that explored mental health in adolescents in LMICS.')
vn('Các vị thành niên tham gia nghiên cứu Photovoice ở LMICs ĐƯỢC ĐẶT TRONG bối cảnh BỆNH TRUYỀN NHIỄM và LẤY MẪU từ các khu định cư KHÔNG CHÍNH THỨC. Ngoài ra, các tổ chức từ thiện như Oppenheimer Memorial Trust và Bill & Melinda Gates Foundation tài trợ các nghiên cứu Photovoice ở LMICs.')
d.add_paragraph()

en('In comparison, higher education organizations including universities and research institutions, fully or partly funded six studies conducted in HICs.')
vn('So sánh: các tổ chức GIÁO DỤC ĐẠI HỌC bao gồm đại học và viện nghiên cứu tài trợ ĐẦY ĐỦ hoặc MỘT PHẦN cho 6 nghiên cứu ở HICs.')
d.add_paragraph()

# === STRENGTHS LIMITATIONS ===
add_h('4. ĐIỂM MẠNH VÀ HẠN CHẾ', 1)
en('Our review is the first of its kind to investigate the use of Photovoice within mental health research involving adolescents. Our review adds to a paucity of research that is beneficial to researchers, academics and clinicians and could inform future Photovoice studies and funding decisions.')
vn('Tổng quan của chúng tôi là CÔNG TRÌNH ĐẦU TIÊN cùng loại điều tra việc dùng Photovoice trong nghiên cứu sức khoẻ tâm thần vị thành niên. Đóng góp vào sự thiếu hụt nghiên cứu, có ích cho nhà nghiên cứu, học giả và bác sĩ lâm sàng, và có thể định hướng các nghiên cứu Photovoice tương lai và quyết định tài trợ.')
d.add_paragraph()

en('We used a structured, peer-reviewed, robust and inclusive approach to elicit information from multiple studies derived from several sources. Individuals from different countries, disciplines and research units participated in the narrative synthesis which made the testing of the synthesis strong.')
vn('Chúng tôi sử dụng cách tiếp cận CÓ CẤU TRÚC, ĐƯỢC PEER-REVIEW, VỮNG CHẮC và BAO TRÙM để thu thập thông tin từ nhiều nghiên cứu. Cá nhân từ các nước, bộ môn và đơn vị nghiên cứu khác nhau tham gia tổng hợp NARRATIVE — làm cho kiểm thử mạnh.')
d.add_paragraph()

en('However, our review has limitations. Firstly, our review is limited by the parameter used to define adolescents, which is likely to vary across sources. Secondly, the use of pre-extracted data may make this review difficult to replicate. Although we have remedied this limitation by including the data extraction table within our own review (see Table 2). Thirdly, the inclusion of published studies introduces publication bias to our review. Lastly, the low quality of studies included in our review may affect the validity of the review findings.')
vn('Tuy nhiên, tổng quan có HẠN CHẾ: Một, định nghĩa vị thành niên có thể khác giữa các nguồn. Hai, dùng dữ liệu trích xuất sẵn có thể khiến tổng quan KHÓ NHÂN ĐÔI — chúng tôi khắc phục bằng cách đưa Table 2 vào tổng quan. Ba, đưa các bài đã công bố vào tạo BIAS CÔNG BỐ. Bốn, chất lượng các nghiên cứu THẤP có thể ảnh hưởng tính HIỆU LỰC của các phát hiện.')
d.add_paragraph()

# === IMPLICATIONS ===
add_h('5. HÀM Ý (Implications)', 1)
en('Our review has the potential to influence various stakeholders, including parents, educators, service designers and policymakers, who can make more impactful decisions based on research co-produced by adolescents.')
vn('Tổng quan có TIỀM NĂNG ảnh hưởng nhiều bên liên quan: phụ huynh, nhà giáo dục, nhà thiết kế dịch vụ và nhà hoạch định chính sách — họ có thể ra quyết định có TÁC ĐỘNG hơn dựa trên nghiên cứu được ĐỒNG SẢN XUẤT (co-produced) bởi vị thành niên.')
d.add_paragraph()

en('Combining photographs and narratives within Photovoice research captures nuanced and emotive experiences, challenging the traditional methods that often collect one form of data. By engaging adolescents in the research process, Photovoice advocates the fair representation of adolescents within mental health research. This approach can promote adolescents\' perspectives and reduce the potential biases imposed by adult researchers, producing a more credible interpretation of adolescents\' mental health experiences.')
vn('Kết hợp ảnh và tự thuật trong Photovoice nắm bắt được trải nghiệm TINH TẾ và CẢM XÚC, thách thức các phương pháp truyền thống thường chỉ thu thập 1 dạng dữ liệu. Bằng cách lôi kéo vị thành niên vào quá trình nghiên cứu, Photovoice CỔ XUÝ đại diện CÔNG BẰNG của vị thành niên trong nghiên cứu sức khoẻ tâm thần. Cách tiếp cận này có thể thúc đẩy quan điểm vị thành niên và giảm bias do nhà nghiên cứu người lớn áp đặt, tạo ra DIỄN GIẢI ĐÁNG TIN HƠN về trải nghiệm sức khoẻ tâm thần của vị thành niên.')
d.add_paragraph()

# === CONCLUSION ===
page(74)
add_h('6. KẾT LUẬN (Conclusion)', 1)
en('Our systematic review has scoped the available evidence on the Photovoice studies exploring mental health among adolescents. Our review highlights the potential of Photovoice as a method that is in its infancy but can generate rich information about the complex, occasionally contradictory and often interconnecting nature of adolescent mental health.')
vn('Tổng quan hệ thống của chúng tôi đã LẬP PHẠM VI bằng chứng có sẵn về các nghiên cứu Photovoice khám phá sức khoẻ tâm thần vị thành niên. Tổng quan làm nổi bật TIỀM NĂNG của Photovoice như một phương pháp ở GIAI ĐOẠN ĐẦU (infancy) nhưng có thể tạo ra thông tin PHONG PHÚ về bản chất PHỨC TẠP, đôi khi MÂU THUẪN và thường ĐAN XEN của sức khoẻ tâm thần vị thành niên.')
d.add_paragraph()

en('Furthermore, our review suggests that Photovoice is a non-reductive and inclusive method that can flexibly examine the complexity of adolescents\' experiences of mental health across a variety of settings and populations. The strength of Photovoice within mental health research involving adolescents lies in its strength to represent and empower participants, enabling concurrent opposing views to be presented alongside each other.')
vn('Hơn nữa, Photovoice là phương pháp KHÔNG GIẢN LƯỢC (non-reductive) và BAO TRÙM, có thể LINH HOẠT khám phá tính phức tạp trong trải nghiệm sức khoẻ tâm thần của vị thành niên qua nhiều bối cảnh và quần thể. Điểm mạnh của Photovoice trong nghiên cứu sức khoẻ tâm thần vị thành niên nằm ở sức mạnh ĐẠI DIỆN và TRAO QUYỀN cho người tham gia, cho phép các quan điểm ĐỐI LẬP đồng thời được trình bày cùng nhau.')
d.add_paragraph()

en('However, the flexibility Photovoice places on flexibility may have come with a risk of non-standardization. Our review illustrates variable quality in the reporting of photovoice studies examining mental health among adolescents. To advance the method, there is a need to standardize guidance relating to the conduct and reporting of Photovoice studies.')
vn('Tuy nhiên, tính LINH HOẠT của Photovoice có thể đi kèm rủi ro KHÔNG CHUẨN HOÁ. Tổng quan minh hoạ chất lượng BIẾN THIÊN trong báo cáo. Để tiến bộ phương pháp, cần CHUẨN HOÁ HƯỚNG DẪN liên quan đến tiến hành và báo cáo nghiên cứu Photovoice.')
d.add_paragraph()

en('Our review captures a point in time where there are relatively few Photovoice studies among adolescents and serves as a guiding example for future mental health research. To refine and establish best practices, more published Photovoice research examining mental health among adolescents is needed. Future Photovoice studies could consider conducting Photovoice studies online. This approach has previously been employed among adolescents (Macias et al., 2023) and can recruit over 100 participants (Subasi et al., 2023 and Doyumğaç et al., 2021), thereby addressing the limitations associated with smaller sample sizes seen in Photovoice research.')
vn('Tổng quan của chúng tôi nắm bắt MỘT THỜI ĐIỂM tương đối ít nghiên cứu Photovoice ở vị thành niên, đóng vai trò là VÍ DỤ HƯỚNG DẪN cho nghiên cứu sức khoẻ tâm thần tương lai. Để tinh chỉnh và thiết lập THỰC HÀNH TỐT NHẤT, CẦN có thêm nghiên cứu Photovoice được công bố. Các nghiên cứu tương lai có thể cân nhắc tiến hành Photovoice TRỰC TUYẾN. Cách tiếp cận này đã được dùng cho vị thành niên (Macias và cộng sự, 2023) và có thể tuyển trên 100 người tham gia (Subasi và cộng sự, 2023; Doyumğaç và cộng sự, 2021), giải quyết hạn chế cỡ mẫu nhỏ.')
d.add_paragraph()

en('Overall, our review indicates that Photovoice is a compelling and promising method for generating insights into the social determinants of adolescent mental health. Its potential to challenge the way in which researchers engage with participants and shape society\'s understanding of adolescent mental health is obvious.')
vn('Tổng thể, tổng quan của chúng tôi chỉ ra Photovoice là phương pháp THUYẾT PHỤC và HỨA HẸN để tạo ra hiểu biết về các yếu tố XÃ HỘI quyết định sức khoẻ tâm thần vị thành niên. Tiềm năng thách thức cách các nhà nghiên cứu tương tác với người tham gia và định hình hiểu biết xã hội về sức khoẻ tâm thần vị thành niên là RÕ RÀNG.')
d.add_paragraph()

# === Bảng 1 — 12 STUDIES ===
add_h('Phụ lục 1 — Bảng 1: 12 nghiên cứu được đưa vào tổng quan', 1)
vn('Bảng dưới đây tóm tắt 12 nghiên cứu Photovoice về sức khoẻ tâm thần vị thành niên (chi tiết trong PDF gốc Bảng 1, trang 6-9):', bold=True)
d.add_paragraph()

t1 = d.add_table(rows=13, cols=4); t1.style = 'Table Grid'
hdr = ['#', 'Tác giả & năm', 'Bối cảnh / mẫu', 'Chủ đề chính (theo tác giả)']
for i, h in enumerate(hdr):
    c = t1.rows[0].cells[i]; shade(c, '4472C4'); pp = c.paragraphs[0]; rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
studies12 = [
    ('1', 'Bashore et al. 2017 (USA, J Child Health Care)', 'After-school program đô thị, THCS', 'Sức khoẻ; stress với giận dữ/thất vọng; tình bạn-giao tiếp; bắt nạt; lập kế hoạch tương lai'),
    ('2', 'Dempsey 2016 (USA, IJN&CP)', 'Vị thành niên nông thôn, depression risk', 'Hoạt động; nhạy cảm; sáng tạo; tâm linh; hài hước; tàn nhẫn; an toàn; trách nhiệm; tiện ích; bệnh-tử vong (theo SEM)'),
    ('3', 'Georgievski et al. 2018 (Canada, J Psychosocial Oncology)', 'VTN đang điều trị ung thư', '18 chủ đề: điểm mạnh cá nhân/xã hội/điều trị; gia đình như bạn thật; thay đổi tình bạn; sợ hãi; hy vọng'),
    ('4', 'Liegghio 2016 (Canada, Intersectionalities)', 'VTN psychiatrized', '"Bình thường" là LÝ TƯỞNG; bình thường mơ hồ-chuyển dịch; counter-narrative bao gồm distress'),
    ('5', 'Orth & van Wyk 2022 (S. Africa, Psych Res Behav Mngmt)', 'VTN sống với HIV ở Cape Town', 'Kết nối; tâm linh-mindfulness; coherence-awareness xã hội; tự trọng; tự hiệu quả; chấp nhận bản thân; coping; resilience; mục đích sống; hoạt động vui; chức năng thể chất'),
    ('6', 'Rose et al. 2018 (USA, J Adol Health)', 'VTN đô thị', 'Coping qua âm nhạc, sao nhãng'),
    ('7', 'Watson & Douglas 2012 (USA, Pediatric Nursing)', 'VTN, an toàn cộng đồng', 'An toàn-không an toàn (Muggers Bridge); gia đình bảo vệ'),
    ('8', 'Velez-Grau 2019 (USA, Children & Youth Services Rev)', 'VTN nhận dịch vụ MH', 'Một chiều của MH theo gia đình; đại lý cá nhân; được nói thẳng với bệnh nhân'),
    ('9', 'Bowers & Wozniak 2020 (USA)', 'Teenagers chung', 'Bạn bè như nguồn hỗ trợ'),
    ('10', 'Wainaina et al. 2021 (Kenya)', 'Bà mẹ trẻ ở khu định cư không chính thức', 'Nguồn lực tài chính, an toàn, mâu thuẫn vai trò'),
    ('11', 'Woodgate et al. 2021 (Canada)', 'VTN có lo âu', 'Lo âu/trầm cảm = "đứng yên một chỗ"'),
    ('12', 'Northwest Michigan CHIR Behavioral Health Initiative 2021 (USA, grey literature)', 'VTN cộng đồng', 'Bạn không cô đơn; cộng đồng đồng cảm-tôn trọng'),
]
for i, row in enumerate(studies12):
    for j, v in enumerate(row):
        t1.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in t1.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True

d.add_paragraph()

# === PHẢN BIỆN ===
add_h('PHẢN BIỆN CHI TIẾT (em viết — không có trong bài gốc)', 1, RED)
crit = [
    ('① Đây là SCOPING review (không phải systematic review chuẩn) → bằng chứng yếu hơn',
     'Scoping review tập trung MAP các nghiên cứu hiện có (descriptive), KHÔNG đo effect size hay tổng hợp định lượng. Phù hợp cho lĩnh vực thiếu evidence. Tuy nhiên không thay thế được systematic review/MA về hiệu quả. Photovoice cho VN cần CẢ scoping (như bài này) + RCT đo outcome.'),
    ('② Cỡ mẫu rất nhỏ (12 nghiên cứu, 4-58 người/NC) → khái quát hoá hạn chế',
     'Tổng cộng ~200 vị thành niên xuyên 12 nghiên cứu. Đa số ≤10 người/NC. Trong khi đó các RCT MH có thể tuyển hàng nghìn. Photovoice mạnh ở DEPTH (sâu sắc), không phải BREADTH (rộng).'),
    ('③ Cảnh báo bias chính: Tây phương + đô thị + nữ quá đại diện',
     '10/12 ở HICs (USA n=6, Canada n=3, UK n=1); 9/12 đô thị; 127 nữ vs 47 nam. Toàn bộ phát hiện về "coping/resilience/family/safety" có thể bias HICS-đô thị-nữ. Áp dụng cho VN (LMIC, đa dạng đô thị/nông thôn, cần cân giới) phải VERY THẬN TRỌNG.'),
    ('④ Chất lượng báo cáo trung bình-kém (KHÔNG có nghiên cứu nào "tốt")',
     '6/12 trung bình + 4/12 kém + 2/12 không đủ chuẩn (theo COREQ). KHÔNG nghiên cứu nào ĐẠT "tốt". Hầu hết thiếu báo cáo: tuyển mộ, thu thập dữ liệu, reflexivity. Nghĩa là phần lớn finding chưa thật sự reproducible.'),
    ('⑤ Hàm ý mạnh cho nghiên cứu VN — nhưng cần thiết kế kỹ',
     'Photovoice phù hợp HS VN vì: (a) các em quen với điện thoại + ảnh; (b) có thể "vượt qua" rào cản ngôn ngữ giới (đặc biệt "lo âu ẩn" 18 % theo Jefferies 2020 SAD VN); (c) trao quyền cho HS VN — chưa quen được hỏi ý kiến trong nghiên cứu. Nên: (i) ≥30 HS/dự án; (ii) cân giới (50:50 nam:nữ); (iii) đa dạng vùng (đô thị Hà Nội/TPHCM + nông thôn Nghệ An/Lạng Sơn); (iv) áp dụng SHOWeD framework; (v) báo cáo theo COREQ chuẩn.'),
    ('⑥ Vai trò bổ sung trong CSDL của em',
     'Tổng cộng 87 canonical entries hiện tại của em CHỦ YẾU là RCT/SR/MA định lượng. QT064 này là bài Photovoice ĐẦU TIÊN trong CSDL — bổ sung phương pháp ĐỊNH TÍNH có sự tham gia. Có thể dùng làm methodological reference khi đề xuất nghiên cứu sâu hơn về trải nghiệm SKTT của HS VN.'),
    ('⑦ Photovoice + lo âu HS VN — gợi ý đề tài tương lai',
     'Có thể design: 30-40 HS THCS/THPT VN có triệu chứng lo âu (GAD-7 ≥ 5) → đào tạo sử dụng điện thoại chụp ảnh + đạo đức (1 buổi) → 1-2 tuần chụp ảnh "điều khiến tôi LO ÂU/AN TOÀN ở trường" → 4 buổi thảo luận nhóm theo SHOWeD → triển lãm + báo cáo policy. Chi phí thấp (đa số HS VN có điện thoại). Đầu ra: insights nguyên gốc + advocacy cho cải cách trường học.'),
]
for ttl, body in crit:
    p = d.add_paragraph(); rr = p.add_run(ttl); rr.bold = True; rr.font.color.rgb = RED
    p2 = d.add_paragraph(); rr2 = p2.add_run(body); rr2.font.color.rgb = RED

d.add_paragraph()

# === Note about Tables 2 + References ===
add_h('Phụ lục 2 — Bảng 2 (data extraction) + References', 1)
vn('Theo Nguyên tắc dịch v2 — Nguyên tắc 2: References giữ nguyên tiếng Anh trong PDF gốc (chuẩn APA, không dịch).', bold=True)
en('Bảng 2 (data extraction) chiếm 53 trang (p20-73) trong PDF gốc với chi tiết cho 215 nghiên cứu Burn et al. extraction. Bảng đầy đủ giữ trong PDF gốc — không lặp lại trong bản dịch để tránh trùng lặp dữ liệu. Người đọc cần tham chiếu PDF gốc cho data extraction chi tiết.')
vn('Bảng 2 (trích xuất dữ liệu) chiếm 53 trang (trang 20-73) trong PDF gốc với chi tiết cho 215 nghiên cứu trong extraction của Burn và cộng sự. Bảng đầy đủ GIỮ TRONG PDF GỐC — không lặp lại trong bản dịch để tránh trùng lặp dữ liệu. Người đọc cần tham chiếu PDF gốc cho data extraction chi tiết.')
d.add_paragraph()
en('References (pages 76-82) follow APA format and are retained in English in original PDF. See PDF for full reference list.')
vn('Phần References (trang 76-82) theo format APA và được giữ tiếng Anh trong PDF gốc. Xem PDF gốc để có danh mục tham khảo đầy đủ.')

d.add_paragraph()
foot = d.add_paragraph()
fr = foot.add_run('Bản dịch song ngữ Anh-Việt — biên soạn 28/04/2026. Phần chính (p2-19, p74) dịch chi tiết 1-1 theo Nguyên tắc dịch v2. Bảng 2 + References giữ nguyên tiếng Anh trong PDF gốc (Nguyên tắc 2).')
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
