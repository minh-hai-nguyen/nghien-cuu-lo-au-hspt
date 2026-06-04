"""Build 00_Khung tap huan-v2.docx — sua + bo sung dua tren:
- Walder 2025 (QT040) DMHI MA 21 RCT g=0,878 SAD-specific
- Qiaochu 2025 (QT050) Mobile CBT 9 RCT
- Samele Clear Fear App 2025 (QT062) 6 modules CBT
- EACP Lochman 2025 (QT065) 25 buoi HS + 16 buoi cha me
- Cai 2025 (Frontiers Psych) 21 RCT school-based resilience
- Peer Support Murphy 2024 (QT066)
- So lieu thuc trang chuong 3 (Bang 3.20-3.45)

CHU XANH (BLUE) cho phan SUA/BO SUNG.
CHU DEN (BLACK) cho phan giu nguyen tu file goc.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/bang-so-lieu-binh-luan/00_Khung tập huấn-v2.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)  # màu xanh cho phần sửa
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')


def H(text, bold=True, size=14, color=NAVY, align=None):
    p = d.add_paragraph()
    if align is not None:
        p.alignment = align
    r = p.add_run(text); r.bold = bold
    r.font.size = Pt(size); r.font.color.rgb = color


def P(text, color=BLACK, size=13, indent=True, bold=False, italic=False):
    """Đoạn văn — màu BLACK = giữ nguyên, BLUE = sửa/bổ sung."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color


def MIX(parts):
    """Tạo đoạn văn có CHẦM xanh-đen.
    parts = [(text, color), ...]
    """
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.0)
    for text, color in parts:
        r = p.add_run(text); r.font.size = Pt(13); r.font.color.rgb = color


def add_table_colored(header, rows, blue_cells=None):
    """Tạo bảng với một số ô tô xanh.
    blue_cells = set of (row_idx, col_idx) — 0-indexed (header = row 0).
    """
    blue_cells = blue_cells or set()
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        cell = tbl.rows[0].cells[i]
        cell.text = ''
        p = cell.paragraphs[0]
        r = p.add_run(h); r.font.size = Pt(10); r.font.name = 'Times New Roman'; r.bold = True
        if (0, i) in blue_cells:
            r.font.color.rgb = BLUE
    for ri, row in enumerate(rows, 1):
        for ci, c in enumerate(row):
            cell = tbl.rows[ri].cells[ci]
            cell.text = ''
            p = cell.paragraphs[0]
            r = p.add_run(str(c)); r.font.size = Pt(10); r.font.name = 'Times New Roman'
            if (ri, ci) in blue_cells:
                r.font.color.rgb = BLUE


# ============================================================
# Title
# ============================================================
H('3.8 Đề xuất Khung Chương trình tập huấn phòng ngừa rối loạn lo âu ở học sinh THCS dựa trên kết quả nghiên cứu',
  size=14, color=NAVY)
P('Bản v2 — sửa và bổ sung dựa trên: Walder và cộng sự (2025) DMHI cho SAD; '
  'Qiaochu & Wang (2025) Mobile CBT; Samele và cộng sự (2025) Clear Fear app; '
  'Lochman và cộng sự (2025) EACP RCT 40 trường; Cai và cộng sự (2025) school-based '
  'resilience; Murphy và cộng sự (2024) peer support; số liệu thực trạng chương 3 '
  'luận án. Phần SỬA/BỔ SUNG được tô MÀU XANH; phần giữ nguyên màu đen.',
  color=BLUE, italic=True, indent=False)

# ============================================================
# 3.8.1 Cơ sở khoa học
# ============================================================
H('3.8.1 Cơ sở khoa học của Chương trình tập huấn phòng ngừa rối loạn lo âu ở học sinh THCS', size=13)

H('3.8.1.1 Cơ sở lý thuyết', size=13)
P('Cơ sở lý thuyết của chương trình tập huấn phòng ngừa rối loạn lo âu ở học sinh THCS '
  'được xây dựng trên sự kết hợp giữa tâm lý học phát triển lứa tuổi THCS, các mô hình '
  'giải thích rối loạn lo âu, cách tiếp cận sức khỏe tâm thần học đường theo toàn trường, '
  'và bằng chứng thực nghiệm từ nghiên cứu thực trạng. Cách kết hợp này nhằm bảo đảm '
  'chương trình không chỉ đúng về mặt lý thuyết mà còn phù hợp với đặc điểm học sinh '
  'THCS và nhu cầu thực tiễn của nhà trường.')

P('- Thứ nhất, chương trình dựa trên cơ sở tâm lý học phát triển lứa tuổi thiếu niên sớm. '
  'Học sinh THCS đang ở giai đoạn có nhiều thay đổi mạnh về sinh học, nhận thức, cảm xúc '
  'và quan hệ xã hội; đồng thời đây cũng là thời kỳ chuyển tiếp học đường, khi yêu cầu '
  'học tập tăng lên và sự nhạy cảm với đánh giá của người khác trở nên rõ hơn. Vì vậy, '
  'rối loạn lo âu ở lứa tuổi này thường gắn chặt với áp lực học tập, hình ảnh bản thân, '
  'quan hệ bạn bè và cảm giác thuộc về trường lớp (Steinberg, 2014). Đây là lý do chương '
  'trình phải được thiết kế theo hướng gần với trải nghiệm đời sống học sinh, nhấn mạnh '
  'kỹ năng tự nhận diện cảm xúc, tự điều chỉnh và tìm kiếm hỗ trợ, thay vì chỉ cung cấp '
  'kiến thức khái niệm.')

# Sửa thứ 2: bổ sung CBT 8-12 buổi
MIX([
    ('- Thứ hai, cơ sở lý thuyết quan trọng nhất của chương trình là mô hình NHẬN THỨC – '
     'HÀNH VI (Cognitive Behavioral Theory – CBT). Theo tiếp cận này, lo âu không chỉ '
     'xuất phát từ sự kiện bên ngoài mà còn được duy trì bởi cách cá nhân diễn giải tình '
     'huống, các suy nghĩ tự động tiêu cực và các hành vi né tránh (Beck, 1976; Hofmann '
     'và cộng sự, 2012). Do đó, một chương trình phòng ngừa lo âu cho học sinh THCS cần '
     'giúp các em nhận ra mối liên hệ giữa nhận thức – cảm xúc – hành vi, nhận diện các '
     'kiểu suy nghĩ sai lệch như sợ sai, sợ bị đánh giá, thảm họa hóa kết quả học tập, '
     'và học cách thay thế bằng cách diễn giải linh hoạt hơn. ', BLACK),
    ('Cụ thể, các CHƯƠNG TRÌNH CBT HỌC ĐƯỜNG hiện đại tại châu Á thường được thiết kế '
     'theo cấu trúc 8–12 BUỔI – mỗi buổi 60–90 phút – tích hợp sáu mô-đun lõi: (1) tâm '
     'lý giáo dục về lo âu (psychoeducation); (2) theo dõi triệu chứng (symptom '
     'tracking); (3) thử thách suy nghĩ thảm họa (cognitive challenges); (4) kỹ thuật '
     'thư giãn (relaxation tools — thở 4-7-8, grounding 5-4-3-2-1); (5) kích hoạt hành '
     'vi (behavioral activation); (6) thang phơi nhiễm (exposure ladder) — phù hợp với '
     'cấu trúc app Clear Fear (Samele và cộng sự, 2025) và mô hình EACP của Lochman và '
     'cộng sự (2025). ', BLUE),
    ('Các nghiên cứu gần đây cho thấy can thiệp học đường dựa trên CBT là một trong những '
     'hướng có bằng chứng tốt nhất đối với lo âu ở trẻ em và thanh thiếu niên (Hofmann '
     'và cộng sự, 2012)', BLACK),
    ('; phân tích tổng hợp 21 RCT của Walder và cộng sự (2025) trên Journal of Medical '
     'Internet Research xác lập thêm: can thiệp sức khỏe tâm thần số (Digital Mental '
     'Health Intervention – DMHI) THIẾT KẾ ĐẶC THÙ cho lo âu xã hội đạt hiệu lực LỚN '
     '(Hedges g = 0,878), gần gấp đôi DMHI tổng quát (g = 0,508). Đây là cơ sở để chương '
     'trình bổ sung thành phần iCBT (internet-delivered CBT) cho học sinh có nguy cơ cao.', BLUE),
])

P('- Thứ ba, chương trình dựa trên khung năng lực học tập cảm xúc – xã hội (Social and '
  'Emotional Learning, SEL), đặc biệt là khung của CASEL. Theo khung này, sự phát triển '
  'của học sinh không chỉ phụ thuộc vào kiến thức học thuật mà còn vào các năng lực như '
  'tự nhận thức, tự quản lý, nhận thức xã hội, kỹ năng quan hệ và ra quyết định có trách '
  'nhiệm. Đối với chương trình phòng ngừa lo âu, khung SEL cung cấp nền tảng để xây dựng '
  'các nội dung như nhận diện cảm xúc, điều hòa cảm xúc, kỹ năng giao tiếp, xử lý xung '
  'đột, tìm kiếm hỗ trợ và duy trì quan hệ tích cực với gia đình, thầy cô, bạn bè (Durlak '
  'và cộng sự, 2011). Điều này đặc biệt quan trọng trong bối cảnh lo âu học đường không '
  'chỉ là vấn đề nội tâm mà còn chịu ảnh hưởng mạnh từ các mối quan hệ trong hệ sinh '
  'thái học đường.')

# Sửa thứ 4: bổ sung whole-school evidence
MIX([
    ('- Thứ tư, chương trình được đặt trong cách tiếp cận sức khỏe tâm thần học đường '
     'theo TOÀN TRƯỜNG. Các nghiên cứu cho thấy phòng ngừa sức khỏe tâm thần trong nhà '
     'trường sẽ hiệu quả hơn khi kết hợp giữa can thiệp cá nhân và môi trường học đường, '
     'bao gồm vai trò của giáo viên, gia đình và văn hóa lớp học (Weare & Nind, 2011; '
     'WHO, 2021). Vì vậy, chương trình không thiết kế như một chuỗi bài giảng rời, mà '
     'là một can thiệp trong hệ sinh thái học đường. ', BLACK),
    ('Phân tích tổng hợp 21 RCT can thiệp resilience tại trường của Cai và cộng sự (2025) '
     'trên Frontiers in Psychiatry xác lập kích thước hiệu ứng SMD = 0,17 (KTC 95% '
     '0,06–0,29; p < 0,01) — hiệu quả NHỎ NHƯNG CÓ Ý NGHĨA THỐNG KÊ. Đáng chú ý, '
     'heterogeneity rất CAO (I² = 81,90%) gợi ý hiệu quả phụ thuộc thiết kế can thiệp – '
     'do đó chương trình cần CỤ THỂ HÓA cho bối cảnh Việt Nam.', BLUE),
])

P('- Thứ năm, chương trình kế thừa cách tiếp cận giáo dục kỹ năng sống và năng lực thích '
  'ứng trong nhà trường. Giáo dục kỹ năng sống được xem là nền tảng giúp học sinh ứng phó '
  'hiệu quả với các thách thức thông qua các kỹ năng như quản lý cảm xúc, giải quyết vấn '
  'đề và đối phó với stress (WHO, 1997). Điều này phù hợp với mục tiêu phòng ngừa lo âu '
  'ở học sinh THCS trong bối cảnh học tập và đời sống số hiện nay.')

P('- Thứ sáu, chương trình còn dựa trên mô hình nguy cơ – bảo vệ. Theo mô hình này, rối '
  'loạn lo âu là kết quả của sự tương tác giữa các yếu tố nguy cơ và yếu tố bảo vệ '
  '(Garmezy, 1991; Masten, 2014). Dữ liệu thực nghiệm cho thấy yếu tố nguy cơ có vai trò '
  'chi phối mạnh, trong khi các yếu tố bảo vệ như tự trọng và hỗ trợ xã hội giúp giảm '
  'thiểu lo âu. Vì vậy, chương trình được thiết kế theo nguyên tắc: giảm nguy cơ và tăng '
  'cường bảo vệ.')

# Sửa thứ 7: bổ sung mô hình cha mẹ EACP
MIX([
    ('- Thứ bảy, chương trình cũng dựa trên cơ sở lý thuyết về sự THAM GIA CỦA GIA ĐÌNH '
     'trong hỗ trợ sức khỏe tâm thần học đường. Các nghiên cứu cho thấy sự hỗ trợ của '
     'cha mẹ có ảnh hưởng đáng kể đến sức khỏe tâm thần của trẻ vị thành niên (Yap và '
     'cộng sự, 2014). Điều này cho thấy cần tích hợp thành phần phối hợp gia đình trong '
     'chương trình. ', BLACK),
    ('Mô hình Early Adolescent Coping Power (EACP) của Lochman và cộng sự (2025) — RCT '
     'cụm 40 trường THCS Mỹ với n = 709 học sinh lớp 7 — gợi ý cấu trúc cụ thể: 25 buổi '
     'cho học sinh + 16 BUỔI 90 PHÚT cho CHA MẸ + thành phần phối hợp với giáo viên. '
     'Phát hiện đặc biệt: nữ học sinh trong nhóm EACP có giảm khó khăn ở trường có ý '
     'nghĩa thống kê — phù hợp với phát hiện chương 3 luận án về chênh lệch giới tính '
     '(F gender × RLLATQ = 44,484; F × RLLAXH = 45,984; cả hai p < 0,001).', BLUE),
])

# Bổ sung thứ 8 về peer support
MIX([
    ('- Thứ tám, chương trình bổ sung thành phần HỖ TRỢ ĐỒNG ĐẲNG (peer support). Scoping '
     'review của Murphy và cộng sự (2024) trong Journal of Community Psychology trên 15 '
     'nghiên cứu, 13 can thiệp peer support cho thấy peer support có TIỀM NĂNG cải thiện '
     'kết quả phục hồi (recovery) — với điều kiện vai trò peer worker được định nghĩa rõ '
     'và có hỗ trợ chuyên môn từ giáo viên/cố vấn học đường.', BLUE),
])

P('Từ các cơ sở lý thuyết trên, có thể khái quát rằng việc phòng ngừa lo âu học đường '
  'hiệu quả cần đồng thời tác động vào cá nhân học sinh, môi trường học đường và hệ '
  'thống hỗ trợ gia đình, qua đó tạo nên sự thay đổi bền vững.')

# ============================================================
# 3.8.1.2 Cơ sở thực tiễn — bổ sung số liệu cụ thể chương 3
# ============================================================
H('3.8.1.2 Cơ sở thực tiễn', size=13)

# Đoạn 13 giữ nguyên
P('Cơ sở thực tiễn của chương trình tập huấn được xây dựng dựa trên kết quả nghiên cứu '
  'thực trạng rối loạn lo âu ở học sinh trung học cơ sở, qua đó cung cấp bằng chứng khoa '
  'học về mức độ biểu hiện và các yếu tố ảnh hưởng trong bối cảnh học đường hiện nay.')

# Sửa: bổ sung số liệu cụ thể
MIX([
    ('Trước hết, kết quả nghiên cứu cho thấy rối loạn lo âu ở học sinh THCS có mức độ '
     'phổ biến ở ngưỡng trung bình, với sự khác biệt giữa các dạng lo âu. ', BLACK),
    ('Cụ thể, trên thang RCADS quy đổi 0–100, ĐTB lo âu lan tỏa (RLLATQ) = 55,82 (cao '
     'nhất); lo âu xã hội (RLLAXH) = 48,41; lo âu chia ly (RLLAC) = 25,06 (thấp nhất). '
     'Nếu áp dụng ngưỡng 65/100 (T-score borderline), ước lượng (giả định normal) tỷ lệ '
     'học sinh có triệu chứng vượt ngưỡng là khoảng 33,85% cho RLLATQ; 25,98% cho RLLAXH; '
     'và 5,01% cho RLLAC – cho thấy quy mô vấn đề ở mức KHẨN CẤP cho lo âu lan tỏa và '
     'xã hội. ', BLUE),
    ('Trong đó, lo âu tổng quát có xu hướng nổi trội hơn so với lo âu xã hội và lo âu '
     'chia ly, phản ánh đặc điểm tâm lý của học sinh trong bối cảnh áp lực học tập và '
     'yêu cầu thích nghi ngày càng cao. Các biểu hiện lo âu tập trung chủ yếu vào áp '
     'lực học tập, sự tự đánh giá bản thân và các tình huống xã hội, cho thấy lo âu '
     'không chỉ mang tính cá nhân mà còn gắn chặt với môi trường học đường và các mối '
     'quan hệ xã hội của học sinh.', BLACK),
])

MIX([
    ('Thứ hai, kết quả phân tích các yếu tố ảnh hưởng chỉ ra rằng rối loạn lo âu chịu '
     'tác động đồng thời của nhiều yếu tố nguy cơ và yếu tố bảo vệ. Cụ thể, các yếu tố '
     'nguy cơ như áp lực học tập, nghiện điện thoại và trải nghiệm bắt nạt học đường có '
     'tác động cùng chiều đến mức độ lo âu, ', BLACK),
    ('với hệ số chuẩn hóa β trong mô hình SEM lần lượt là: ÁP LỰC HỌC TẬP β = 0,510 '
     '(RLLATQ) và β = 0,490 (RLLAXH) – CƯỜNG ĐỘ LỚN NHẤT; NGHIỆN ĐIỆN THOẠI β = 0,336 '
     'và 0,383; BẮT NẠT HỌC ĐƯỜNG β = 0,215 (RLLATQ), nhưng đặc biệt MẠNH NHẤT cho lo '
     'âu CHIA LY (β RLLACL = 0,376). Yếu tố nguy cơ TỔNG HỢP có β = 0,747 và giải thích '
     '55,8% phương sai rối loạn lo âu (R² = 0,558). ', BLUE),
    ('trong đó áp lực học tập nổi lên như một yếu tố trung tâm trong bối cảnh giáo dục '
     'hiện nay. Điều này phản ánh thực tế rằng học sinh không chỉ chịu áp lực từ yêu '
     'cầu học tập mà còn từ sự cạnh tranh và kỳ vọng xã hội. Bên cạnh đó, việc sử dụng '
     'điện thoại thông minh quá mức và các trải nghiệm tiêu cực trong quan hệ bạn bè '
     'góp phần làm gia tăng nguy cơ lo âu ở học sinh.', BLACK),
])

MIX([
    ('Ngược lại, các yếu tố bảo vệ như mức độ gắn bó với trường học, sự hỗ trợ xã hội từ '
     'gia đình và bạn bè, cũng như lòng tự trọng có vai trò làm giảm mức độ rối loạn lo '
     'âu. ', BLACK),
    ('Hệ số β chuẩn hóa: TỰ TRỌNG β = −0,455 (RLLATQ) và −0,415 (RLLAXH) – cường độ '
     'BẢO VỆ MẠNH NHẤT; HỖ TRỢ TỪ CHA MẸ β = −0,172 và −0,273 (mạnh nhất cho lo âu xã '
     'hội); GẮN BÓ TRƯỜNG HỌC β = −0,108 và −0,187. Đáng chú ý, HỖ TRỢ TỪ BẠN BÈ GẦN '
     'NHƯ KHÔNG có tác động bảo vệ (β = −0,015 đến −0,079; p > 0,05 cho hai trong ba '
     'dạng RLLA) – TRÁI với y văn quốc tế, gợi ý đặc thù lứa tuổi THCS Việt Nam. Yếu tố '
     'bảo vệ TỔNG HỢP có β = −0,352 và giải thích 12,4% phương sai (R² = 0,124) – yếu '
     'hơn yếu tố nguy cơ tổng (R² = 0,558) gấp 4,5 lần. ', BLUE),
    ('Điều này cho thấy môi trường học đường tích cực và các nguồn lực tâm lý – xã hội '
     'đóng vai trò quan trọng trong việc hỗ trợ học sinh thích nghi và giảm thiểu các '
     'vấn đề tâm lý. Tuy nhiên, mức độ ảnh hưởng của các yếu tố bảo vệ còn hạn chế, cho '
     'thấy các nguồn lực này chưa được khai thác và phát huy một cách hiệu quả trong '
     'thực tiễn giáo dục.', BLACK),
])

MIX([
    ('Đáng chú ý, kết quả nghiên cứu cũng cho thấy các biện pháp đối phó của học sinh '
     'có xu hướng liên hệ cùng chiều với rối loạn lo âu, gợi ý rằng các chiến lược đối '
     'phó hiện tại có thể chưa phù hợp hoặc chưa hiệu quả trong việc kiểm soát căng '
     'thẳng và lo âu. ', BLACK),
    ('Cụ thể, β BPĐP → RLLATQ = 0,749 – cường độ CỰC LỚN, lớn nhất trong toàn bộ chương '
     '3 (R² = 0,561 ≈ 56,1%). Đây là hiện tượng MALADAPTIVE COPING ESCALATION (Compas '
     'và cộng sự, 2017): học sinh càng lo âu càng dùng nhiều biện pháp đối phó, nhưng '
     'nếu không hiệu quả thì lo âu vẫn duy trì hoặc tăng. Tuy nhiên, mô hình SEM cho '
     'BPĐP có FIT INDICES KÉM (RMSEA 0,080–0,204; CFI 0,865–0,911) — phát hiện này nên '
     'được trình bày như EXPLORATORY và cần phân tách Brief-COPE 14 nhân tố theo Carver '
     '(1997) trong nghiên cứu tiếp theo. ', BLUE),
    ('Điều này đặt ra yêu cầu cấp thiết phải trang bị cho học sinh các kỹ năng đối phó '
     'mang tính thích ứng, có định hướng và dựa trên bằng chứng khoa học.', BLACK),
])

MIX([
    ('Ngoài ra, sự khác biệt có ý nghĩa thống kê về mức độ rối loạn lo âu theo giới '
     'tính và khối lớp cho thấy yếu tố phát triển và đặc điểm cá nhân cũng cần được '
     'xem xét trong quá trình thiết kế chương trình can thiệp. ', BLACK),
    ('Cụ thể: F (giới × RLLATQ) = 44,484; F × RLLAXH = 45,984 (cả hai p < 0,001) — '
     'NỮ > NAM; nhưng F × RLLAC = 0,246 (p = 0,620) — KHÔNG có chênh lệch giới ở lo âu '
     'chia ly. Theo Cohen d, hiệu ứng kích thước cho RLLATQ = 0,365 và RLLAXH = 0,370 '
     '— GẦN NHƯ BẰNG NHAU, đều thuộc nhóm small-to-medium. Đối với khối lớp, RLLAC '
     'GIẢM ĐƠN ĐIỆU theo lớp (32,13 → 27,14 → 20,88 → 19,46 từ lớp 6 đến lớp 9) – phù '
     'hợp khung phát triển: lo âu chia ly khởi phát sớm và giảm khi tự lập. ', BLUE),
    ('Điều này hàm ý rằng các nội dung tập huấn cần có sự linh hoạt, phù hợp với từng '
     'nhóm đối tượng cụ thể thay vì áp dụng một cách đồng nhất.', BLACK),
])

P('Tổng hợp các kết quả trên cho thấy rối loạn lo âu ở học sinh THCS là một vấn đề có '
  'tính hệ thống, chịu ảnh hưởng đa chiều từ cá nhân, gia đình và môi trường học đường. '
  'Do đó, việc xây dựng chương trình tập huấn phòng ngừa không thể chỉ tập trung vào một '
  'yếu tố riêng lẻ, mà cần tiếp cận theo hướng tích hợp, kết hợp giữa giảm thiểu yếu tố '
  'nguy cơ, tăng cường yếu tố bảo vệ và phát triển các kỹ năng đối phó hiệu quả cho học '
  'sinh. Đây chính là cơ sở thực tiễn quan trọng để đề xuất và thiết kế chương trình tập '
  'huấn phòng ngừa rối loạn lo âu trong bối cảnh trường trung học cơ sở hiện nay.')

# ============================================================
# 3.8.2 Mục tiêu, nguyên tắc, đối tượng
# ============================================================
H('3.8.2 Mục tiêu, nguyên tắc, đối tượng và phạm vi áp dụng của chương trình tập huấn '
  'phòng ngừa rối loạn lo âu cho học sinh trung học cơ sở', size=13)

H('3.8.2.1 Mục tiêu', size=13)
MIX([
    ('Trong bối cảnh sức khỏe tâm thần học sinh ngày càng được quan tâm, các chương '
     'trình phòng ngừa rối loạn lo âu trong trường học được thiết kế nhằm can thiệp sớm '
     'vào các yếu tố nguy cơ và tăng cường các yếu tố bảo vệ tâm lý. Các nghiên cứu gần '
     'đây cho thấy các chương trình tập huấn kiến thức dựa trên trường học có thể giúp '
     'giảm đáng kể mức độ lo âu và cải thiện khả năng thích ứng của học sinh, đặc biệt '
     'khi tập trung vào phát triển kỹ năng nhận thức, điều chỉnh cảm xúc (Stockings và '
     'cộng sự, 2022; Dunning và cộng sự, 2022). ', BLACK),
    ('Trên cơ sở đó, mục tiêu của chương trình tập huấn phòng ngừa rối loạn lo âu ở học '
     'sinh trung học cơ sở được CỤ THỂ HÓA thành các mục tiêu cốt lõi: (1) GIẢM YẾU TỐ '
     'NGUY CƠ (áp lực học tập, nghiện điện thoại, bắt nạt học đường) – mục tiêu giảm '
     '|β| > 0,3 trong mô hình SEM hậu can thiệp; (2) TĂNG YẾU TỐ BẢO VỆ (lòng tự trọng, '
     'hỗ trợ cha mẹ, gắn bó trường học) – đặc biệt nhắm vào tự trọng do là yếu tố bảo '
     'vệ MẠNH NHẤT (β = −0,455); (3) PHÁT TRIỂN KỸ NĂNG ĐỐI PHÓ THÍCH ỨNG để giảm '
     'maladaptive coping escalation; (4) TÍCH HỢP THÀNH PHẦN DIGITAL (iCBT app) cho '
     'nhóm có nguy cơ cao, dựa trên Walder và cộng sự (2025). Hiệu quả được đánh giá '
     'qua so sánh trước-sau bằng các thang chuẩn quốc tế (RCADS, GAD-7, ESSA, RSES, '
     'MSPSS, PSSM).', BLUE),
])

H('3.8.2.2 Nguyên tắc', size=13)
P('Chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh trung học cơ sở được xây '
  'dựng dựa trên nguyên tắc sau:')
P('- Chương trình được xây dựng dựa trên bằng chứng khoa học, kế thừa các kết quả nghiên '
  'cứu thực nghiệm và các mô hình can thiệp đã được kiểm chứng, nhằm đảm bảo tính hiệu '
  'quả và độ tin cậy của nội dung can thiệp (Stockings và cộng sự, 2022).')
P('- Chương trình đảm bảo tính phù hợp văn hóa, được điều chỉnh theo bối cảnh giáo dục '
  'và đặc điểm văn hóa – xã hội của học sinh Việt Nam, nhằm tăng khả năng tiếp nhận và '
  'hiệu quả thực tiễn.')
P('- Nội dung chương trình phù hợp với đặc điểm phát triển của lứa tuổi trung học cơ sở, '
  'chú trọng đến đặc điểm sinh lý, nhận thức và cảm xúc của học sinh trong giai đoạn vị '
  'thành niên, từ đó lựa chọn phương pháp và nội dung can thiệp phù hợp.')
P('- Chương trình được thiết kế theo tiếp cận đa tầng, kết hợp giữa phòng ngừa phổ quát '
  'và phòng ngừa chọn lọc, nhằm tác động đồng thời đến toàn bộ học sinh và nhóm có nguy '
  'cơ cao, qua đó nâng cao hiệu quả phòng ngừa một cách toàn diện (WHO, 2021).')
# Thêm nguyên tắc 5 — digital component
MIX([
    ('- ', BLACK),
    ('Bổ sung nguyên tắc thứ năm: chương trình tích hợp THÀNH PHẦN DIGITAL (iCBT app) '
     'cho nhóm có nguy cơ cao, dựa trên bằng chứng phân tích tổng hợp 21 RCT của Walder '
     'và cộng sự (2025) cho thấy DMHI thiết kế đặc thù SAD đạt Hedges g = 0,878. App '
     'phải có HƯỚNG DẪN NGƯỜI (g = 0,825 vs g ≈ 0,2 không hướng dẫn) – có thể qua '
     'chat-bot hoặc tin nhắn từ giáo viên/cố vấn học đường. CẢNH BÁO: Mobile CBT có '
     'hiệu quả YẾU cho lo âu khi không thiết kế đặc thù (Qiaochu & Wang, 2025: chỉ 2/6 '
     'RCT có ý nghĩa) – do đó app phải SAD-specific hoặc có thành phần CBT đầy đủ '
     'theo cấu trúc Clear Fear (Samele và cộng sự, 2025).', BLUE),
])

H('3.8.2.3 Đối tượng và phạm vi áp dụng', size=13)
P('- Đối tượng chính của chương trình là học sinh THCS, nhóm lứa tuổi đang trải qua '
  'những biến đổi mạnh mẽ về sinh lý, nhận thức và cảm xúc, đồng thời có nguy cơ cao '
  'xuất hiện các rối loạn lo âu nếu không được hỗ trợ kịp thời.')
P('- Đối tượng phối hợp bao gồm giáo viên, phụ huynh và cán bộ tư vấn học đường, là '
  'những lực lượng có vai trò trực tiếp trong việc phát hiện, hỗ trợ và duy trì hiệu '
  'quả can thiệp trong môi trường tự nhiên của học sinh, góp phần tăng tính bền vững của '
  'chương trình (Fazel và cộng sự, 2014).')
P('- Phạm vi áp dụng của chương trình được triển khai theo hướng phòng ngừa phổ quát đối '
  'với toàn bộ học sinh nhằm nâng cao nhận thức và kỹ năng tâm lý cơ bản, đồng thời kết '
  'hợp phòng ngừa chọn lọc đối với nhóm học sinh có nguy cơ cao, nhằm tối ưu hóa hiệu '
  'quả can thiệp và sử dụng nguồn lực hợp lý (WHO, 2021).')

# ============================================================
# 3.8.3 Khung chương trình
# ============================================================
H('3.8.3 Khung chương trình tập huấn phòng ngừa rối loạn lo âu ở học sinh THCS', size=13)

H('3.8.3.1 Cấu trúc chương trình', size=13)
MIX([
    ('Bảng dưới đây trình bày cấu trúc chương trình tập huấn phòng ngừa rối loạn lo âu '
     'học đường dành cho học sinh THCS. ', BLACK),
    ('Cấu trúc gồm 8 chủ đề, tổng cộng 19 BUỔI (8 lý thuyết + 11 thực hành), tương đương '
     '~28–32 GIỜ – phù hợp với mô hình can thiệp CBT 8–12 buổi của Walder và cộng sự '
     '(2025) và EACP của Lochman và cộng sự (2025) đã sử dụng tại châu Á và Mỹ. Bảng có '
     'bổ sung cột THAM CHIẾU KHOA HỌC và MODULE DIGITAL để làm rõ cơ sở của từng chủ đề.', BLUE),
])

# Bảng cấu trúc — bổ sung 2 cột mới (xanh)
add_table_colored(
    ['TT', 'Nội dung tập huấn', 'Số buổi (LT/TH)', 'Thời lượng',
     'Mục tiêu', 'Tham chiếu khoa học', 'Module digital'],
    [
        ['1', 'Nhận diện rối loạn lo âu học đường', '1 LT + 1 TH', '90–120 phút',
         'Giúp HS hiểu bản chất và biểu hiện lo âu',
         'Beck (1976); Hofmann (2012)', 'Module 1 Clear Fear: psychoeducation'],
        ['2', 'Áp lực học tập và lo âu', '1 LT + 1 TH', '90–120 phút',
         'Nhận diện nguồn áp lực học tập',
         'Sun (2011) ESSA; Pascoe (2020); β=0,510',
         'Module 2 Clear Fear: symptom tracking'],
        ['3', 'Kỹ năng quản lý áp lực học tập', '1 LT + 2 TH', '135–180 phút',
         'Trang bị kỹ năng giảm áp lực',
         'Walder (2025) g=0,508',
         'Module 5 Clear Fear: behavioral activation'],
        ['4', 'Kiểm soát sử dụng điện thoại', '1 LT + 1 TH', '90–120 phút',
         'Giảm nguy cơ từ hành vi số',
         'Schmidt-Persson (2024) RCT; β=0,336',
         'App theo dõi screen time'],
        ['5', 'Ứng phó với bắt nạt học đường', '1 LT + 1 TH', '90–120 phút',
         'Tăng kỹ năng tự bảo vệ',
         'Olweus (1996); β=0,376 cho RLLACL',
         'OBPP module trực tuyến'],
        ['6', 'Xây dựng tự trọng tích cực', '1 LT + 2 TH', '135–180 phút',
         'Tăng yếu tố bảo vệ nội tại',
         'Rosenberg (1965); β=−0,455 mạnh nhất',
         '—'],
        ['7', 'Điều chỉnh suy nghĩ & kỹ năng đối phó', '1 LT + 2 TH', '135–180 phút',
         'Giảm suy nghĩ tiêu cực và né tránh',
         'Beck CBT; Carver (1997) Brief-COPE',
         'Module 3+4+6 Clear Fear: cognitive challenges, relaxation, exposure ladder'],
        ['8', 'Kết nối hỗ trợ gia đình – nhà trường', '1 LT + 1 TH', '90–120 phút',
         'Tăng nguồn lực hỗ trợ',
         'Yap (2014); EACP 16 buổi cha mẹ; Murphy (2024) peer support',
         'Phụ huynh: chat-bot hỗ trợ'],
    ],
    blue_cells={
        # Header xanh cho 2 cột mới
        (0, 5), (0, 6),
        # Cell tham chiếu (col 5) cho rows 1-8
        (1, 5), (2, 5), (3, 5), (4, 5), (5, 5), (6, 5), (7, 5), (8, 5),
        # Cell module digital (col 6) cho rows 1-8
        (1, 6), (2, 6), (3, 6), (4, 6), (5, 6), (6, 6), (7, 6), (8, 6),
    }
)

P('')

# ============================================================
# 3.8.3.2 Nội dung chi tiết — chỉ note ngắn vì nội dung dài đã có
# ============================================================
H('3.8.3.2 Nội dung chương trình', size=13)
MIX([
    ('Nội dung chi tiết 8 chủ đề được giữ nguyên từ bản gốc. Phần BỔ SUNG ở mỗi chủ đề – '
     'tích hợp thành phần CBT cụ thể và module digital – được tô màu xanh trong bảng cấu '
     'trúc trên. Đặc biệt:', BLUE),
])
MIX([
    ('• ', BLACK),
    ('Chủ đề 7 (Điều chỉnh suy nghĩ và kỹ năng đối phó) là TRỤ CỘT của CBT 8–12 buổi: '
     'lý thuyết giới thiệu mô hình ABC (Activating event – Belief – Consequence) của '
     'Beck; thực hành theo cấu trúc 6 module Clear Fear: cognitive challenges (thử thách '
     'suy nghĩ thảm họa – ví dụ "Nếu thi không đạt thì cuộc đời tôi sẽ kết thúc" → '
     '"Nếu thi không đạt thì có thể thi lại lần sau"); relaxation tools (thở 4-7-8, '
     'grounding 5-4-3-2-1); exposure ladder (leo thang phơi nhiễm – ví dụ HS sợ phát '
     'biểu trước lớp: bước 1 – nói trước gương; bước 2 – nói trước 1 bạn thân; bước 3 – '
     'nói trước nhóm 3 bạn; ...). ', BLUE),
])
MIX([
    ('• ', BLACK),
    ('Chủ đề 8 (Kết nối hỗ trợ gia đình – nhà trường) cần cấu trúc phù hợp với mô hình '
     'EACP của Lochman và cộng sự (2025): 16 BUỔI 90 PHÚT cho cha mẹ với các nội dung '
     'cốt lõi – (a) nhận diện dấu hiệu lo âu ở con; (b) kỹ năng giao tiếp không phán xét; '
     '(c) tránh thiên lệch "kỳ vọng cao" làm tăng áp lực học tập; (d) phối hợp với nhà '
     'trường khi con có dấu hiệu lo âu mức borderline.', BLUE),
])
MIX([
    ('• ', BLACK),
    ('Đối với HS có nguy cơ CAO (T-score RCADS ≥ 65), bổ sung iCBT APP với 6 module '
     'theo Clear Fear (Samele và cộng sự, 2025) – có hỗ trợ chat-bot hoặc giáo viên/cố '
     'vấn (theo Walder 2025: guided DMHI g = 0,825 vs unguided g ≈ 0,2). Thời lượng app '
     '8–9 tuần, tương ứng với 8–12 buổi tập huấn nhóm.', BLUE),
])

P('Nội dung 8 chủ đề chi tiết (1) Nhận diện rối loạn lo âu học đường; (2) Áp lực học '
  'tập và tác động đến rối loạn lo âu; (3) Kỹ năng quản lý áp lực học tập; (4) Kiểm soát '
  'sử dụng điện thoại; (5) Ứng phó với bắt nạt học đường; (6) Xây dựng tự trọng tích '
  'cực; (7) Điều chỉnh suy nghĩ và kỹ năng đối phó với lo âu; (8) Tăng cường hỗ trợ từ '
  'gia đình và nhà trường — được giữ nguyên từ bản gốc với ba điểm bổ sung trên.')

# ============================================================
# 3.8.4 Đánh giá hiệu quả
# ============================================================
H('3.8.4 Đánh giá hiệu quả chương trình tập huấn phòng ngừa rối loạn lo âu học đường '
  'dành cho học sinh THCS', size=13)
P('Việc đánh giá hiệu quả chương trình được thực hiện theo thiết kế đa tầng, bao gồm cấp '
  'độ toàn trường, nhóm nhỏ và cá nhân, nhằm xác định mức độ thay đổi của học sinh sau '
  'tập huấn trên ba phương diện: mức độ rối loạn lo âu, các yếu tố nguy cơ và các yếu tố '
  'bảo vệ.')
MIX([
    ('- Ở cấp độ toàn trường, đánh giá được tiến hành theo mô hình đánh giá trước – sau '
     'tập huấn nhằm xác định hiệu quả phòng ngừa phổ quát. Các chỉ số chính bao gồm '
     'điểm trung bình rối loạn lo âu (tổng quát, xã hội, chia ly), mức độ các yếu tố '
     'nguy cơ (áp lực học tập, nghiện điện thoại, bắt nạt học đường) và yếu tố bảo vệ '
     '(tự trọng, hỗ trợ từ cha mẹ, gắn bó với trường học). Hiệu quả được xác định khi '
     'các chỉ số rối loạn lo âu và nguy cơ giảm, trong khi các yếu tố bảo vệ tăng có ý '
     'nghĩa thống kê. ', BLACK),
    ('Đề xuất sử dụng các thang chuẩn quốc tế: RCADS (Chorpita, 2000) cho RLLA; ESSA 16 '
     'mục (Sun, 2011) cho áp lực học tập; SAS-SV hoặc BSMAS cho nghiện điện thoại; '
     'OBVQ-R (Olweus, 1996) cho bắt nạt; RSES (Rosenberg, 1965) cho tự trọng; MSPSS '
     '(Zimet, 1988) cho hỗ trợ xã hội; PSSM (Goodenow, 1993) cho gắn bó trường học. '
     'Mức cải thiện có ý nghĩa thống kê với Cohen d ≥ 0,3 (small-to-medium) được xem là '
     'thành công bước đầu; d ≥ 0,5 (medium) được xem là thành công đáng kể.', BLUE),
])
P('- Ở cấp độ nhóm nhỏ, đánh giá tập trung vào học sinh có nguy cơ cao, nhằm xem xét '
  'hiệu quả phòng ngừa trên nhóm chọn lọc. Ngoài các thang đo định lượng, việc đánh giá '
  'được bổ sung bằng quan sát, phản hồi của giáo viên và phụ huynh, cũng như mức độ cải '
  'thiện các kỹ năng cụ thể (quản lý học tập, kiểm soát điện thoại, ứng phó với rối '
  'loạn lo âu).')
P('- Ở cấp độ cá nhân, việc đánh giá nhằm theo dõi sự thay đổi của từng học sinh, đặc '
  'biệt là những trường hợp có biểu hiện rối loạn lo âu nổi bật. Các chỉ báo bao gồm '
  'thay đổi điểm số, hành vi thực tế và mức độ đạt được mục tiêu cá nhân. Dữ liệu được '
  'thu thập thông qua tự đánh giá của học sinh, nhận xét của giáo viên và phụ huynh, '
  'cùng với nhật ký cảm xúc.')
MIX([
    ('Phân tích dữ liệu được thực hiện bằng các phương pháp thống kê như so sánh trước '
     'sau về mức độ và biểu hiện. Chương trình được xem là hiệu quả khi mức độ rối loạn '
     'lo âu giảm, yếu tố nguy cơ suy giảm, yếu tố bảo vệ gia tăng và học sinh thể hiện '
     'sự cải thiện rõ rệt về kỹ năng và hành vi. ', BLACK),
    ('Khuyến nghị sử dụng SEM trước-sau (latent change SEM) để kiểm tra mô hình tích '
     'hợp YTNC + YTBV → RLLA có thay đổi sau can thiệp; đặc biệt mục tiêu kéo β YTNC '
     'từ 0,747 xuống dưới 0,5 và đẩy |β YTBV| từ 0,352 lên trên 0,4 sau can thiệp.', BLUE),
])
P('Như vậy, cách tiếp cận đánh giá đa cấp độ này cho phép phản ánh đầy đủ hiệu quả của '
  'chương trình ở cả mức độ phổ quát, chọn lọc và cá nhân, đồng thời cung cấp cơ sở '
  'khoa học cho việc điều chỉnh và triển khai rộng rãi trong môi trường học đường.')

# ============================================================
# 3.8.5 Phương pháp và quy trình triển khai
# ============================================================
H('3.8.5 Phương pháp và quy trình triển khai chương trình tập huấn phòng ngừa rối loạn '
  'lo âu cho học sinh THCS', size=13)

H('3.8.5.1 Phương pháp triển khai', size=13)
MIX([
    ('Chương trình tập huấn phòng ngừa rối loạn lo âu học đường cho học sinh THCS được '
     'triển khai theo tiếp cận đa tầng (phổ quát – chọn lọc – cá nhân hóa), nhằm bảo '
     'đảm tác động đồng thời đến toàn bộ học sinh và các nhóm có nguy cơ cao. Về phương '
     'pháp, chương trình được tổ chức dưới hình thức tập huấn theo nhóm lớp, lồng ghép '
     'vào sinh hoạt lớp hoặc hoạt động trải nghiệm, sử dụng cách tiếp cận lấy học sinh '
     'làm trung tâm. Trong đó, học sinh không chỉ tiếp nhận kiến thức mà còn tham gia '
     'các hoạt động thực hành như thảo luận tình huống, đóng vai, tự đánh giá và xây '
     'dựng kế hoạch cá nhân. Tỷ lệ giữa lý thuyết và thực hành được duy trì cân đối nhằm '
     'tăng khả năng vận dụng vào thực tế. Đối với học sinh có nguy cơ cao, nhà trường '
     'tổ chức các hoạt động hỗ trợ bổ sung theo nhóm nhỏ hoặc cá nhân, đồng thời phối '
     'hợp với phụ huynh và giáo viên để tăng hiệu quả can thiệp. ', BLACK),
    ('Bổ sung kênh DIGITAL: học sinh có nguy cơ cao (T-score ≥ 65 trên RCADS) được giới '
     'thiệu sử dụng iCBT app dạng Clear Fear với 6 module theo cấu trúc CBT chuẩn (xem '
     'mục 3.8.3.1). App phải có thành phần GUIDED – chat-bot hoặc tin nhắn định kỳ từ '
     'giáo viên/cố vấn học đường để tăng hiệu lực (Walder, 2025: guided g = 0,825). '
     'Đối với SAD (lo âu xã hội) – dạng RLLA có chênh lệch giới rõ rệt – nên dùng app '
     'SAD-SPECIFIC (g = 0,878 theo Walder) thay vì app tổng quát.', BLUE),
])

H('3.8.5.2 Quy trình triển khai', size=13)
P('Quy trình triển khai chương trình được tổ chức theo các giai đoạn liên tục và có hệ '
  'thống, bao gồm:')
P('-\tGiai đoạn chuẩn bị: Xây dựng kế hoạch triển khai, tập huấn giáo viên, chuẩn bị '
  'tài liệu và công cụ đánh giá, đồng thời thông tin và phối hợp với phụ huynh.')
P('-\tGiai đoạn đánh giá ban đầu: Khảo sát mức độ lo âu, các yếu tố nguy cơ và yếu tố '
  'bảo vệ nhằm xác định thực trạng và nhóm học sinh cần ưu tiên hỗ trợ.')
P('-\tGiai đoạn triển khai tập huấn: Tổ chức các buổi tập huấn theo trình tự logic từ '
  'nhận diện rối loạn lo âu, giảm yếu tố nguy cơ (áp lực học tập, nghiện điện thoại, '
  'bắt nạt học đường), đến tăng yếu tố bảo vệ (tự trọng, kỹ năng đối phó) và tăng cường '
  'nguồn lực hỗ trợ (gia đình, nhà trường).')
P('-\tGiai đoạn đánh giá sau tập huấn: Tiến hành đo lường lại các chỉ số đã khảo sát ban '
  'đầu để xác định mức độ thay đổi và hiệu quả của chương trình.')
P('-\tGiai đoạn theo dõi và duy trì: Tiếp tục hỗ trợ học sinh có nguy cơ, tổ chức các '
  'hoạt động củng cố và duy trì môi trường học đường tích cực.')

MIX([
    ('Về thời điểm triển khai, chương trình nên được thực hiện vào đầu năm học, khi học '
     'sinh đang trong giai đoạn thích nghi và áp lực học tập chưa gia tăng mạnh, nhằm '
     'đạt hiệu quả phòng ngừa tối ưu. Ngoài ra, có thể triển khai vào đầu học kỳ II để '
     'can thiệp kịp thời khi áp lực bắt đầu tăng. Cần tránh triển khai vào các giai '
     'đoạn cao điểm kiểm tra và thi cử để bảo đảm sự tham gia đầy đủ và hiệu quả của '
     'học sinh. ', BLACK),
    ('Về thời lượng tổng: 8 chủ đề × ~2,5 buổi trung bình = 20 buổi × 90–120 phút = '
     '~30–40 GIỜ. Đối với cha mẹ: 16 BUỔI 90 phút theo mô hình EACP (Lochman, 2025). '
     'Đối với HS có nguy cơ cao: thêm 8–9 TUẦN sử dụng iCBT app. ', BLUE),
    ('Như vậy, phương pháp và quy trình triển khai chương trình được thiết kế theo hướng '
     'khoa học, hệ thống và khả thi, phù hợp với điều kiện thực tiễn của trường THCS '
     'và góp phần nâng cao hiệu quả phòng ngừa rối loạn lo âu học đường.', BLACK),
])

# ============================================================
# Phụ lục TLTK bổ sung
# ============================================================
H('Phụ lục — Tài liệu tham khảo BỔ SUNG (chỉ liệt kê các tài liệu MỚI dùng để nâng cấp '
  'lên v2)', size=13, color=BLUE)

new_refs = [
    'Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, 1594658. https://doi.org/10.3389/fpsyt.2025.1594658',
    'Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100.',
    'Compas, B. E., và cộng sự. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991.',
    'Goodenow, C. (1993). The Psychological Sense of School Membership Scale (PSSM). Psychology in the Schools, 30(1), 79–90.',
    'Lochman, J. E., và cộng sự. (2025). Randomized controlled trial of the Early Adolescent Coping Power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology. [QT065]',
    'Murphy, R., và cộng sự. (2024). A systematic scoping review of peer support interventions in integrated primary youth mental health care. Journal of Community Psychology, 52(1), 154–180. [QT066]',
    'Olweus, D. (1996). The Revised Olweus Bully/Victim Questionnaire (OBVQ-R). University of Bergen.',
    'Qiaochu, Z., & Wang, Y. (2025). Effectiveness of Mobile-Based Cognitive Behavioural Therapy for Depression and Anxiety in Children and Young People: A Systematic Review of RCTs. Clinical Psychology & Psychotherapy, 32(6), e70173. [QT050]',
    'Samele, C., và cộng sự. (2025). Clear Fear app for adolescents and young adults: A pre-post evaluation. JMIR Formative Research. [QT062]',
    'Schmidt-Persson, J., và cộng sự. (2024). Screen media use and mental health of children and adolescents: A secondary analysis of a randomized clinical trial. JAMA Network Open, 7(7), e2419881. [QT033]',
    'Sun, J., Dunne, M. P., Hou, X. Y., & Xu, A. Q. (2011). Educational Stress Scale for Adolescents (ESSA). Journal of Psychoeducational Assessment, 29(6), 534–546.',
    'Walder, N., và cộng sự. (2025). Digital mental health interventions for prevention and treatment of social anxiety disorder in children, adolescents and young adults: Systematic review and meta-analysis. Journal of Medical Internet Research, JMIR. [QT040]',
    'Zimet, G. D., và cộng sự. (1988). The Multidimensional Scale of Perceived Social Support (MSPSS). Journal of Personality Assessment, 52(1), 30–41.',
]
for r in new_refs:
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    run = p.add_run(r); run.font.size = Pt(11); run.font.color.rgb = BLUE

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
