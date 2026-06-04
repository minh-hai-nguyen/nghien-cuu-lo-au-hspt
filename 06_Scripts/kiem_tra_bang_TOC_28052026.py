# -*- coding: utf-8 -*-
"""Inspect manual TOC table - read structure, find correct table.
28/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

# Find MỤC LỤC paragraph + the table after it
toc_para_idx = None
for i, p in enumerate(d.paragraphs[:200]):
    if p.text.strip().upper() in ('MỤC LỤC', 'MUC LUC'):
        toc_para_idx = i
        break
print(f"MỤC LỤC heading at paragraph {toc_para_idx}")

# Find tables — locate the TOC table (table containing 'MỞ ĐẦU' as first entry)
print()
print("=== All tables - find TOC table ===")
toc_table_idx = None
for ti, tb in enumerate(d.tables):
    if len(tb.rows) > 0 and len(tb.rows[0].cells) >= 2:
        first_cell = tb.rows[0].cells[0].text.strip()
        if first_cell.upper() == 'MỞ ĐẦU' or 'Lý do' in first_cell:
            toc_table_idx = ti
            print(f"  Table {ti}: {len(tb.rows)} rows, first cell: {first_cell!r}")
            break

if toc_table_idx is None:
    # Search for TOC table by content
    for ti, tb in enumerate(d.tables):
        if len(tb.rows) >= 3:
            # Check first few rows for "Lý do" or "MỞ ĐẦU"
            for ri in range(min(5, len(tb.rows))):
                cell0 = tb.rows[ri].cells[0].text.strip()
                if 'MỞ ĐẦU' in cell0.upper() or 'Lý do chọn đề tài' in cell0:
                    toc_table_idx = ti
                    print(f"  Found at Table {ti}, row {ri}: {cell0!r}")
                    break
            if toc_table_idx is not None: break

if toc_table_idx is None:
    print("  KHONG TIM THAY bang TOC!")
else:
    print()
    print(f"=== Bang TOC = Table {toc_table_idx} ===")
    tb = d.tables[toc_table_idx]
    print(f"  Rows: {len(tb.rows)} | Cols: {len(tb.columns)}")
    print()
    print(f"  Content (first 30 rows):")
    for ri, row in enumerate(tb.rows):
        cells_text = [c.text.strip() for c in row.cells]
        if ri < 30:
            print(f"    Row {ri}: {cells_text}")
