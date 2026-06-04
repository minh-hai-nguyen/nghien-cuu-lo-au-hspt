# -*- coding: utf-8 -*-
"""Sinh bo cuc chi tiet IMRaD cho 2 bai Q1 + Q3 theo phuong an A.
- Q1: Integrated SEM model (predictive)
- Q3: Descriptive cross-sectional (item-level + demographic)
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'BoCuc_Q1_Q3_PhuongAnA_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

s = d.styles['Normal']
s.font.name = 'Times New Roman'
s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.3


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    return p


def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    return p


def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    return p


def P(text, italic=False, indent=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.italic = italic
    return p


def B(text, level=0):
    """Bullet"""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('• ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    return p


# ============================================================
# COVER
# ============================================================
H1('BỐ CỤC CHI TIẾT 2 BÀI BÁO Q1 + Q3')
P('(Phương án A — Split theo phương pháp + câu hỏi nghiên cứu)',
  italic=True)
P('NCS Công Thị Hằng — TS. Đào Minh Đức',
  italic=True)
P('Ngày soạn: 01/06/2026', italic=True)
P('')
P('Cơ sở dữ liệu: 1.352 HS THCS Hà Nội (614 nam / 738 nữ) — '
  'từ file thầy gửi "Thực trạng RLLA HS THCS"', italic=True)

d.add_page_break()


# ============================================================
# PART 1: BÀI Q1
# ============================================================
H1('PHẦN 1 — BÀI BÁO Q1')

H2('1.1 Thông tin chung')
B('Tên đề tài (tentative): "Integrated risk-protective structural equation model of '
  'anxiety disorder subtypes among Vietnamese lower secondary school students: '
  'A mixed-methods study"')
B('Target journal (priority order): (1) BMC Psychiatry (Q1, IF ~4.4); '
  '(2) Journal of Affective Disorders (Q1, IF ~6.6); '
  '(3) Asian Journal of Psychiatry (Q1, IF ~12)')
B('Word count: 6.000–8.000 (excluding references)')
B('Authors byline: Hang Thi Cong (1st) + Duc Minh Dao (corresponding) + '
  '1-2 đồng tác giả TBD')


H2('1.2 Novel contributions (3 điểm mới)')
B('(1) **First integrated SEM** kết hợp đồng thời 3 yếu tố nguy cơ (academic '
  'pressure, smartphone addiction, school bullying) + 4 yếu tố bảo vệ (school '
  'engagement, parental support, peer support, self-esteem) → 3 dạng RLLA theo '
  'DSM-5 trên mẫu HS THCS Việt Nam')
B('(2) **Phát hiện gender-invariance** của separation anxiety (F=0,246; p=0,620) '
  'trái với GAD (F=44,484; p<0,001) và Social Anxiety (F=45,984; p<0,001) — '
  'thách thức giả định gender-difference đồng đều cho mọi anxiety subtype')
B('(3) **Mixed-methods integration**: dữ liệu định tính phỏng vấn (HS NND lớp 9, '
  'MNA lớp 9) làm phong phú interpretation của các β coefficients — đặc biệt '
  'cho yếu tố bảo vệ "parental support" và "peer support"')


H2('1.3 Cấu trúc IMRaD')

H3('Title + Abstract (~250 từ — structured Abstract chuẩn BMC Psychiatry)')
B('Background: anxiety prevalence in adolescents + Vietnam gap + integrated model rationale')
B('Methods: 1,352 lower secondary students; 8 validated scales; CFA → SEM → '
  'multi-group invariance')
B('Results: SEM fit indices; β coefficients risk + protective; gender invariance '
  'finding; key qualitative themes')
B('Conclusions: prevention implications; cultural context; methodological '
  'contribution to Vietnam literature')

H3('Introduction (~1.200 từ — 4 paragraphs)')
B('§1 Burden of anxiety disorders in adolescents (Asian + global prevalence, '
  'developmental impact) — cite Xu 2021, Anderson 2025, GBD 2025')
B('§2 Risk-protective framework + Vietnam gap — cite Lazarus & Folkman 1984 '
  '(theoretical); Compas 2017 (meta-analysis); systematic review of Vietnam '
  'literature showing methodological gap')
B('§3 Differential gender patterns hypothesis — cite McLean 2011; Wen 2020; '
  'V-NAMHS 2022 — point out conflicting evidence for separation anxiety')
B('§4 Present study + 3 hypotheses: H1 risk factors positive on all 3 subtypes; '
  'H2 protective factors negative on all 3 subtypes; H3 gender invariance for '
  'separation anxiety only')

H3('Methods (~1.500 từ)')
B('2.1 Participants & sampling: 1.352 HS từ 2 trường THCS Hà Nội (Nhật Tân '
  '+ Tây Mỗ); 614 nam / 738 nữ; tuổi 11-14 (khối 6-9); proportional sampling')
B('2.2 Measures (8 validated scales):', 0)
B('Anxiety: RCADS (Chorpita, 2000) — 3-factor adaptation (GAD 7 items, SAD 4, '
  'SocAD 4)', 1)
B('Risk: ESSA (Sun et al., 2011, 4 items), SAS-SV (Kwon et al., 2013, 5 items), '
  'OBVQ (Olweus, 1996, 8 items)', 1)
B('Protective: PSSM (Goodenow, 1993, 7 items), MSPSS (Zimet et al., 1988, 8 '
  'items, 2 subscales), RSES (Rosenberg, 1965, 5 items)', 1)
B('Score conversion: All raw scores converted to 0-100 scale for cross-scale '
  'comparability', 1)
B('2.3 Qualitative interviews: semi-structured, n=12 purposeful sample (high/'
  'medium/low anxiety subgroups), 30-45 minutes each, thematic analysis '
  '(Braun & Clarke, 2006)')
B('2.4 Analytic strategy:', 0)
B('Descriptive statistics (M, SD) + Cronbach\'s α + McDonald\'s ω', 1)
B('CFA per scale (fit: CFI≥0.90, RMSEA≤0.08)', 1)
B('Integrated SEM model: 7 latent predictors → 3 anxiety latent outcomes', 1)
B('Multi-group invariance by gender (configural → metric → scalar) — focus '
  'separation anxiety', 1)
B('Mixed-methods integration: joint display matrix (Creswell & Plano Clark, '
  '2018)', 1)
B('2.5 Ethics: informed consent (parents + students); ethics approval HNUE')

H3('Results (~1.800 từ)')
B('3.1 Sample characteristics: demographic table (1.352 students)')
B('3.2 Psychometric properties: CFA results per scale (Cronbach\'s α + ω + '
  'fit indices)')
B('3.3 Descriptive anxiety levels: 3 subtypes M, SD')
B('3.4 SEM main model:', 0)
B('Risk paths: ALHT β=+0,533*** > NĐT β=+0,400*** > BNHĐ β=+0,276***', 1)
B('Protective paths: TTr β=-0,457*** > HTCM β=-0,273*** > GBTH β=-0,108*** > '
  'HTBB β=-0,015 ns', 1)
B('R²: total anxiety = 0,284 (risk) + 0,209 (protective) = ~50% explained', 1)
B('3.5 Multi-group invariance by gender — KEY FINDING:', 0)
B('GAD: female > male (Δχ² significant)', 1)
B('Social: female > male', 1)
B('**Separation: gender-invariant (F=0,246; p=0,620)** — full invariance', 1)
B('3.6 Mixed-methods integration: 3-4 qualitative themes confirming/extending '
  'quantitative findings (e.g., parental support → "không chia sẻ" theme from '
  'MNA lớp 9; school engagement → community belonging)')

H3('Discussion (~1.700 từ)')
B('4.1 Summary of key findings (3 hypotheses, gender invariance)')
B('4.2 Risk factors interpretation: academic pressure cultural specificity '
  '(Vietnamese exam-driven education) — compare Xu 2021 China')
B('4.3 Self-esteem as strongest protective lever (~85-89% of academic pressure '
  'magnitude per LA Conclusions)')
B('4.4 Separation anxiety gender-invariance — NOVEL FINDING:', 0)
B('Possible explanations: developmental task (early adolescence pre-pubertal); '
  'cultural collectivism (Vietnamese family attachment cross-gender)', 1)
B('Compare with limited Western evidence (need to search PubMed thoroughly)', 1)
B('4.5 Mixed-methods value: qualitative quotes illuminating quantitative '
  'β coefficients')
B('4.6 Clinical/educational implications: tiered prevention (universal vs '
  'targeted); Khung CT 8 nội dung từ LA')
B('4.7 Limitations: cross-sectional design; 2 schools (Hanoi only); self-report '
  'bias; qualitative sub-sample size')
B('4.8 Future directions: longitudinal design; multi-site VN; intervention RCT')

H3('References (~40-50 refs)')
B('Verified PDFs from `02_Papers-goc/`: Xu 2021, Chen 2023, Wen 2020, Anderson 2025, '
  'Compas 2017, Pascoe 2020, Steare 2023, Jefferies 2020, Nakie 2022, '
  'V-NAMHS 2022, Hoang Trung Hoc 2025, Bhardwaj 2020, Qiu 2022, Zhu 2025, '
  'Alharbi 2019, Saikia 2023')
B('Theoretical: Lazarus & Folkman 1984, Carver 1997, Rosenberg 1965, '
  'Goodenow 1993, Zimet 1988, Chorpita 2000, Sun 2011, Kwon 2013, Olweus 1996, '
  'Hu & Bentler 1999, McLean 2011, Bronfenbrenner & Morris 2006')
B('Method: Braun & Clarke 2006, Creswell & Plano Clark 2018')

H3('Tables + Figures plan')
B('Table 1: Sample demographics')
B('Table 2: Psychometric properties (α, ω, CFA fit) — 8 scales')
B('Table 3: SEM β coefficients (risk + protective → 3 anxiety subtypes)')
B('Table 4: Multi-group invariance fit indices (configural/metric/scalar)')
B('Table 5: Qualitative themes joint display')
B('Figure 1: Integrated SEM model diagram')
B('Figure 2: Multi-group gender comparison plot')

d.add_page_break()


# ============================================================
# PART 2: BÀI Q3
# ============================================================
H1('PHẦN 2 — BÀI BÁO Q3')

H2('2.1 Thông tin chung')
B('Tên đề tài (tentative): "Manifestations and patterns of anxiety disorder subtypes '
  'among Vietnamese lower secondary school students: A cross-sectional descriptive study"')
B('Target journal (priority): (1) BMC Pediatrics; (2) Vietnam Journal of Psychology '
  '(English version); (3) Vietnam Journal of Educational Sciences; (4) PLOS ONE')
B('Word count: 3.500–5.000 (Q3 standard)')
B('Authors byline: Hang Thi Cong (1st) + Duc Minh Dao (corresponding) + '
  '1 đồng tác giả')


H2('2.2 Novel contributions (3 điểm mới — phù hợp Q3 tier)')
B('(1) **First detailed item-level analysis** của RCADS Vietnamese adaptation '
  '(rút gọn 21→15 items) trên mẫu HS THCS lớn N=1.352 — cung cấp normative data')
B('(2) **Cross-sectional grade trajectory** (khối 6→7→8→9) cho 3 dạng RLLA — '
  'phát hiện **separation anxiety giảm liên tục** (32,13 → 27,14 → 20,88 → 19,46) '
  'phù hợp với DSM-5 developmental literature')
B('(3) **Item-level identification of high-anxiety expressions** — top items như '
  '"lo lắng khi nghĩ rằng mình đã không làm tốt điều gì đó" (M=64,28) — cung cấp '
  'mục tiêu cụ thể cho screening + targeted intervention ở VN')


H2('2.3 Cấu trúc IMRaD (đơn giản hơn Q1)')

H3('Title + Abstract (~200 từ)')
B('Background: anxiety in adolescents + Vietnam under-researched item-level')
B('Methods: 1,352 lower secondary students; RCADS Vietnamese; descriptive + ANOVA')
B('Results: M/SD per item; demographic differences; grade trajectory')
B('Conclusions: implications for screening + prevention curriculum')

H3('Introduction (~800 từ — 3 paragraphs)')
B('§1 Adolescent anxiety burden + Asian context — Xu 2021, Wen 2020, Saikia 2023')
B('§2 Vietnam gap: limited item-level data; need normative data + grade '
  'trajectory — V-NAMHS 2022, Hoang Trung Hoc 2025')
B('§3 Present study: descriptive analysis of RCADS Vietnamese adaptation; '
  '3 research questions (RQ1 item-level distribution per subtype; RQ2 gender + '
  'grade differences; RQ3 high-anxiety items as intervention targets)')

H3('Methods (~1.000 từ)')
B('2.1 Participants & sampling: 1.352 HS THCS Hà Nội (giống Q1 — same dataset)')
B('2.2 Instrument: RCADS Vietnamese adaptation (Chorpita 2000 original + '
  'translation/adaptation process); 15 items in 3 subtypes (GAD 7 + SAD 4 + '
  'SocAD 4); 4-point Likert (converted to 0-100)')
B('2.3 Reliability check: Cronbach\'s α + McDonald\'s ω per subtype (already '
  'reported but Q3 needs basic psychometric)')
B('2.4 Analytic strategy:', 0)
B('Descriptive statistics: M, SD, range per item', 1)
B('Ranking items within each subtype (1=highest M to lowest)', 1)
B('Gender comparison: independent t-test or ANOVA per subtype', 1)
B('Grade comparison: one-way ANOVA + post-hoc Tukey', 1)
B('Effect sizes: Cohen\'s d (gender) + η² (grade)', 1)
B('2.5 Ethics: identical to Q1 (consent + HNUE approval)')

H3('Results (~1.200 từ)')
B('3.1 Sample characteristics (brief table)')
B('3.2 Reliability: α + ω per subtype', 0)
B('3.3 Generalized anxiety items (Table from Bảng 1 thầy gửi):', 0)
B('Top item: "lo lắng khi nghĩ rằng mình đã không làm tốt điều gì đó" M=64,28', 1)
B('Lowest item: "lo lắng về những gì sẽ xảy ra" M=45,86', 1)
B('Subtype mean: M=55,82', 1)
B('3.4 Separation anxiety items (Table from Bảng 2):', 0)
B('Top: "sẽ cảm thấy sợ nếu phải ở xa nhà qua đêm" M=27,88', 1)
B('Bottom: "lo lắng khi đi ngủ vào ban đêm" M=21,52', 1)
B('Subtype mean: M=25,06 (lowest of 3 subtypes)', 1)
B('3.5 Social anxiety items (Table from Bảng 3):', 0)
B('Top: "lo lắng về việc người khác nghĩ gì về mình" M=56,98', 1)
B('Bottom: "lo rằng mình sẽ trông ngốc nghếch" M=42,09', 1)
B('Subtype mean: M=48,41', 1)
B('3.6 Demographic comparisons (from Bảng 4):', 0)
B('Gender: female > male for GAD (59,47 vs 51,43; F=44,484; p<0,001) and '
  'Social (52,74 vs 43,20; F=45,984; p<0,001); **separation NS** (F=0,246; '
  'p=0,620)', 1)
B('Grade trajectory: GAD increases 6→9 (54,32 to 59,79); separation '
  'DECREASES 6→9 (32,13 to 19,46) — developmental pattern', 1)

H3('Discussion (~1.000 từ — 4 paragraphs)')
B('4.1 Summary of item-level findings — "academic worry" dominant in GAD; '
  '"social judgment" dominant in social anxiety')
B('4.2 Developmental pattern: separation anxiety declines (consistent with '
  'DSM-5 separation anxiety = childhood-onset); GAD increases (consistent with '
  'late-childhood/adolescence onset literature)')
B('4.3 Gender pattern: female > male except separation anxiety — needs '
  'replication; brief mention not deep dive (reserve for Q1)')
B('4.4 Implications for screening: top items per subtype as screener questions; '
  'limitations (cross-sectional, 2 schools); future replication')

H3('References (~25-35 refs)')
B('Verified PDFs: Xu 2021, Wen 2020, Anderson 2025, V-NAMHS 2022, '
  'Hoang Trung Hoc 2025, Bhardwaj 2020, Nakie 2022, Saikia 2023')
B('Theoretical: Chorpita 2000 (RCADS), Rosenberg 1965 (norm theory)')

H3('Tables + Figures plan (đơn giản hơn Q1)')
B('Table 1: Sample demographics')
B('Table 2: Item-level GAD (7 items: M, SD, rank)')
B('Table 3: Item-level Separation (4 items)')
B('Table 4: Item-level Social (4 items)')
B('Table 5: Demographic comparisons (gender + grade)')
B('Figure 1 (optional): Grade trajectory plot for 3 subtypes')


d.add_page_break()


# ============================================================
# PART 3: DATA SPLIT TABLE - ANTI-OVERLAP
# ============================================================
H1('PHẦN 3 — BẢNG PHÂN CHIA DỮ LIỆU CHỐNG TRÙNG LẶP')

H2('3.1 Nguyên tắc anti-overlap')
P('Để TRÁNH self-plagiarism + bị reject "redundant publication", 2 bài chia '
  'dữ liệu theo 4 tiêu chí ortogonal:', italic=True)
B('(a) **Câu hỏi nghiên cứu khác nhau** — Q1 hỏi cơ chế; Q3 hỏi mô tả')
B('(b) **Statistical lens khác nhau** — Q1 dùng SEM (β); Q3 dùng descriptive (M, ranking, ANOVA)')
B('(c) **Outcome khác nhau** — Q1 focus mô hình tổng thể + invariance; '
  'Q3 focus item-level + demographic')
B('(d) **Method paradigm khác nhau** — Q1 mixed-methods; Q3 quantitative-only descriptive')

H2('3.2 Bảng phân bổ dữ liệu')

t = d.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = 'Dữ liệu'
hdr[1].text = 'Bài Q1 (SEM model)'
hdr[2].text = 'Bài Q3 (Descriptive)'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

data_split = [
    ('Bảng 1 (RLLA tổng quát item-level)',
     'KHÔNG dùng — chỉ summary',
     '✓ CHỦ LỰC — full item table'),
    ('Bảng 2 (Separation anxiety item-level)',
     'KHÔNG dùng — chỉ summary',
     '✓ CHỦ LỰC — full item table'),
    ('Bảng 3 (Social anxiety item-level)',
     'KHÔNG dùng — chỉ summary',
     '✓ CHỦ LỰC — full item table'),
    ('Bảng 4 (Demographic comparison)',
     '✓ Background only; focus invariance multi-group',
     '✓ Full table — primary analysis'),
    ('Bảng 5 (3 YTNC item-level)',
     '✓ Aggregate scores → SEM predictors',
     'KHÔNG dùng (out of scope)'),
    ('Bảng 6 (4 YTBV item-level)',
     '✓ Aggregate scores → SEM predictors',
     'KHÔNG dùng (out of scope)'),
    ('SEM β coefficients',
     '✓ PRIMARY FINDING',
     'KHÔNG dùng'),
    ('Multi-group invariance',
     '✓ NOVEL FINDING (separation anxiety)',
     'Mention briefly nhưng KHÔNG dùng technical detail'),
    ('Qualitative interviews (NND, MNA)',
     '✓ Mixed-methods integration',
     'KHÔNG dùng (purely quantitative paper)'),
    ('CFA per scale',
     '✓ All 8 scales',
     'Brief — only RCADS reliability'),
]
for row_data in data_split:
    row = t.add_row().cells
    for i, cell_text in enumerate(row_data):
        row[i].text = cell_text
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if '✓' in cell_text:
                    if 'KHÔNG' not in cell_text:
                        r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                elif 'KHÔNG' in cell_text:
                    r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


d.add_paragraph()

H2('3.3 Common methodological elements (acknowledged in both papers)')
P('Một số thông tin nền tảng buộc xuất hiện ở cả 2 bài — đây là acceptable:',
  italic=True)
B('Same dataset: 1,352 HS THCS Hà Nội (cite cùng dataset, cùng IRB)')
B('Same RCADS instrument (Chorpita 2000) — different analysis scope')
B('Same sample demographics (brief table cả 2 bài)')
B('Cross-reference: Q3 paper cite Q1 paper khi xuất bản ("for the integrated '
  'SEM model see [Cong & Dao, 202X]")')


d.add_page_break()


# ============================================================
# PART 4: TIMELINE + NEXT STEPS
# ============================================================
H1('PHẦN 4 — TIMELINE + BƯỚC TIẾP THEO')

H2('4.1 Timeline đề xuất')

t2 = d.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
hdr2 = t2.rows[0].cells
hdr2[0].text = 'Tuần'
hdr2[1].text = 'Việc Q1'
hdr2[2].text = 'Việc Q3'
for c in hdr2:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

timeline = [
    ('Tuần 1 (01-07/06)',
     'NCS + thầy confirm phương án A; em viết outline chi tiết Q1',
     '(chưa bắt đầu — sau Q1)'),
    ('Tuần 2 (08-14/06)',
     'Em viết Methods + Results',
     'Em viết outline Q3 đồng thời'),
    ('Tuần 3 (15-21/06)',
     'Em viết Discussion + Introduction; NCS review',
     'Em viết Methods + Results Q3'),
    ('Tuần 4 (22-28/06)',
     'Verify fact + plagiarism check + format BMC Psychiatry',
     'Em viết Discussion + Intro Q3'),
    ('Tuần 5 (29/06-05/07)',
     'Submit Q1 → cover letter + suggested reviewers',
     'Verify + plagiarism check Q3'),
    ('Tuần 6 (06-12/07)',
     'Respond peer review nếu có',
     'Submit Q3 → cover letter'),
]
for row_data in timeline:
    row = t2.add_row().cells
    for i, cell_text in enumerate(row_data):
        row[i].text = cell_text
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)


H2('4.2 Câu hỏi cần thầy + NCS confirm trước khi em viết outline chi tiết')
B('(1) **Authorship order Q1 vs Q3** — có giống nhau không (Hang Thi Cong 1st cả 2 bài)?')
B('(2) **Đồng tác giả** — thầy có muốn mời chuyên gia thống kê (cho Q1) hoặc '
  'chuyên gia adolescent psych (cho Q3)?')
B('(3) **Target journal priority** — Q1 ưu tiên BMC Psychiatry hay J Affective '
  'Disorders? (em đề xuất BMC Psychiatry vì IF 4.4 ổn + acceptance rate ~30%)')
B('(4) **Q3 venue** — submit Vietnamese journal (Vietnam J Psychology EN version) '
  'hay quốc tế Q3 thuần (BMC Pediatrics)?')
B('(5) **Ethics approval** — có IRB letter chính thức từ HNUE chưa? (cần kèm '
  'khi submit Q1)')
B('(6) **Qualitative interviews** — số lượng + status: bao nhiêu HS đã phỏng vấn, '
  'transcripts đã có chưa?')


H2('4.3 Cảnh báo academic integrity')
B('Tuyệt đối KHÔNG đụng bài Q2.5 KHGDVN của NCS (Bai1, Bai2) — '
  'memory đã ghi rõ')
B('KHÔNG copy paraphrase từ LA chính — 2 bài viết hoàn toàn original tiếng Anh, '
  'chỉ dùng raw data + statistical results từ LA')
B('Plagiarism check tự kiểm: tránh chuỗi ≥7 từ giống bất kỳ paper nào đã đăng '
  '(Turnitin threshold)')
B('Citation policy: chỉ cite verified PDF trong `02_Papers-goc/` + DOI verifiable')
B('Số liệu KHÔNG được lệch khỏi LA: 1.352 HS, F=0,246 (separation gender), '
  '~85-89% (TTr magnitude), 8 nội dung Khung CT')


# ============================================================
# Clean metadata
# ============================================================
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
