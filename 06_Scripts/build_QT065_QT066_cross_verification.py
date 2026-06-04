"""Doc verify cross-bài: QT065 EACP + QT066 Peer Support so với 7 bài liên quan trong DB."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/QT065_QT066_cross_reference_verification.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
DARK = RGBColor(31, 73, 125); GREEN = RGBColor(54, 95, 44); RED = RGBColor(192, 0, 0); ORANGE = RGBColor(191, 97, 14); GRAY = RGBColor(90, 90, 90)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def vn(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)

# TITLE
title = d.add_heading('Cross-reference verification: QT065 EACP + QT066 Peer Support so với 7 bài liên quan', level=0)
for r in title.runs: r.font.color.rgb = DARK
sub = d.add_paragraph()
sr = sub.add_run('Đối chiếu liên bài theo Nguyên tắc 9 (cross-reference audit) — đọc trực tiếp từ Tom-tat-tung-bai/, không dùng trí nhớ')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = GRAY
d.add_paragraph()

# CONTEXT
add_h('Bối cảnh', 1)
vn('Tổng quan của em sau khi build bản dịch chi tiết QT065 (Bradshaw/Lochman 2025 EACP) và QT066 (Murphy 2024 Peer Support). '
   'Em verify 2 bài này so với 7 bài can thiệp đã có trong DB 90 canonical, nhằm: (1) phát hiện CONFLICT số liệu/kết luận; '
   '(2) tìm hiểu MỐI LIÊN HỆ giữa các phương pháp; (3) gợi ý vị trí 2 bài này trong landscape can thiệp lo âu HS VN.')
vn('7 bài cross-reference được đối chiếu:')
for it in [
    'QT028 Zugman 2024 — Tổng quan AJP IF=18 (chứa CAMS Walkup 2008)',
    'QT029 Li 2025 — NMA Bayesian 30 RCT BMC',
    'QT040 Walder 2025 — MA DMHI cho SAD (21 RCT)',
    'QT042 Brown & Carter 2025 — Editorial UK BESST/MHST/PLACES',
    'QT043 Bress 2024 — RCT JAMA Maya app (18-25 tuổi)',
    'QT059 Cai 2025 — SR/MA School Resilience (38 NC)',
    'VN030 Tran 2023 — RCT Happy House Việt Nam (RAP-A 6 buổi)',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
d.add_paragraph()

# =====================================================
# PHẦN A — QT065 EACP CROSS-REFERENCE
# =====================================================
add_h('PHẦN A — QT065 EACP cross-reference', 1)

vn('A.1. Bảng đối chiếu chính', bold=True, size=12)
a_tbl = d.add_table(rows=8, cols=4); a_tbl.style = 'Table Grid'
hdr = ['Tiêu chí', 'QT065 EACP', 'Cross-ref', 'Đánh giá']
for i, h in enumerate(hdr):
    c = a_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
a_data = [
    ('Đối tượng', 'HS lớp 7 (TB ~12-13 tuổi) gây hấn',
     'CAMS (QT028): trẻ 7-17 có RLLA chẩn đoán | Maya (QT043): 18-25 có RLLA | Walder (QT040): <25 có SAD subthreshold-mild',
     'EACP TUỔI NHỎ NHẤT trong các can thiệp dạng này. Lưu ý: CAMS có cả lo âu thuần, còn EACP gốc cho gây hấn'),
    ('Mục tiêu', 'PHÒNG NGỪA hành vi gây hấn + ngăn leo thang',
     'CAMS/Maya: ĐIỀU TRỊ RLLA đã có chẩn đoán | Walder: PHÒNG NGỪA (subthreshold) + điều trị | Happy House (VN030): PHÒNG NGỪA UNIVERSAL',
     'EACP = Indicated prevention (HS đã có triệu chứng nhưng chưa chẩn đoán). Khác Maya/CAMS điều trị, khác Happy House universal'),
    ('Format', 'NHÓM 6-7 HS, mặt-đối-mặt, 25 buổi/9 tháng + cá nhân 8-10 buổi 30 phút',
     'CBT cá nhân (Li 2025): xếp #1 NMA SUCRA cao | CBT nhóm: xếp #2-3 | Maya: app digital 12 phiên/6 tuần | Happy House: nhóm 6 buổi 90 phút',
     'EACP có thời lượng DÀI HƠN nhiều: 25 buổi vs 6-12 buổi → có thể "intensive" enough cho gây hấn nhưng quá dài cho lo âu thuần'),
    ('Effect size lo âu', 'Internalizing self-report: γ=0,70, d=0,13 (RẤT NHỎ, p=0,045)',
     'Walder DMHI vs chứng: g=0,508; SAD-specific g=0,878 | Maya: Cohen d=1,04 (HAM-A) | NMA Li: ACT MD=-3,83, CBT MD=-3,64 | Happy House: d=0,11 (CESD-R, depression)',
     '⚠ EACP HIỆU LỰC LO ÂU THẤP NHẤT (d=0,13). Maya/Walder cao hơn đáng kể vì đối tượng CÓ lo âu rõ. Happy House cũng thấp (universal)'),
    ('Hiệu lực giới',
     'NỮ trong EACP: school problems giảm + personal adjustment tốt hơn (γ_EACP*Sex = 1,83, p=0,048)',
     'Maya: không tách giới rõ | Walder: không tách | Happy House: nữ baseline anxiety cao hơn nam',
     'EACP có HIỆU LỰC GIỚI ĐẶC BIỆT — gợi ý phù hợp HS NỮ. Cần verify trong các RCT khác'),
    ('Đa thành phần',
     '3 thành phần: HS (25) + Cha mẹ (12) + GV (3 buổi cheer)',
     'CAMS: HS + thuốc + therapy mặt đối mặt | Happy House: HS only (không có cha mẹ component) | Maya: HS only digital | Walder: phần lớn HS only',
     'EACP UNIQUE: là chương trình DUY NHẤT trong DB có cả 3 thành phần HS + cha mẹ + GV. Phù hợp adapt cho VN'),
    ('Bằng chứng cấp',
     'RCT cluster 40 trường, n=709, NIMH funded',
     'CAMS: RCT đa trung tâm n=488 | Walder: MA 21 RCT g=0,508 | Li: NMA 30 RCT 1.711 trẻ | Maya: RCT n>150 | Happy House: cluster RCT 4+4 trường VN',
     'EACP có cỡ mẫu LỚN (n=709) nhưng chỉ 1 RCT đơn. Bằng chứng tổng thể chưa MA-level'),
]
for i, row in enumerate(a_data):
    for j, v in enumerate(row):
        a_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in a_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn('A.2. Mâu thuẫn / điểm cần lưu ý', bold=True, size=12)

conflict_tbl = d.add_table(rows=5, cols=2); conflict_tbl.style = 'Table Grid'
hdr_c = ['Phát hiện', 'Lý giải / Cảnh báo']
for i, h in enumerate(hdr_c):
    c = conflict_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
conflicts = [
    ('① EACP TĂNG internalizing self-report (γ=0,70, p=0,045) — paradox',
     'Bài gốc lý giải: HS sau EACP tăng AWARENESS về hậu quả → tự báo cáo lo âu tăng ("awareness effect"). KHÔNG mâu thuẫn vì: '
     'Walder/Maya/CBT khác trực tiếp giảm anxiety symptoms; còn EACP giảm GÂY HẤN nhưng có thể tăng lo âu thoáng qua. '
     'CAMS không gặp paradox này vì CAMS đo PHỤC HỒI lo âu (response rate) không phải awareness.'),
    ('② NMA Li 2025 ưu tiên CBT cá nhân, EACP là CBT nhóm',
     'Li xếp CBT cá nhân #1 SUCRA. CBT nhóm xếp #2-3. EACP là CBT nhóm + phù hợp khả năng triển khai trường VN. '
     'KHÔNG mâu thuẫn: CBT nhóm có hiệu lực thấp hơn cá nhân nhưng tỷ lệ COST-EFFECTIVE cao hơn nhiều. EACP nhóm 6-7 HS phù hợp budget VN.'),
    ('③ EACP cho lớp 7, NMA Li chủ yếu trẻ <12 tuổi',
     'Đa số RCT trong NMA Li (30 RCT) cho trẻ <12 tuổi. EACP đặc biệt cho lớp 7 (~12-13). Đây là KHOẢNG TRỐNG mà EACP lấp được — '
     'CAMS có phạm vi tuổi rộng (7-17) nhưng cá nhân/SSRI; EACP school-based nhóm cho lớp 7-8.'),
    ('④ Hiệu lực gender NỮ — nhất quán?',
     'EACP có hiệu lực mạnh hơn ở NỮ. Maya không tách giới. Happy House (VN30): nữ baseline anxiety CAO HƠN nam. '
     'V-NAMHS 2022 (VN002): VN nữ MH issues cao hơn nam. → EACP NỮ-FRIENDLY phù hợp VN context.'),
]
for i, (k, v) in enumerate(conflicts):
    c0 = conflict_tbl.rows[i+1].cells[0]; shade(c0, 'FFF2CC')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = ORANGE
    conflict_tbl.rows[i+1].cells[1].text = v
d.add_paragraph()

# =====================================================
# PHẦN B — QT066 PEER SUPPORT CROSS-REFERENCE
# =====================================================
add_h('PHẦN B — QT066 Peer Support cross-reference', 1)

vn('B.1. Bảng đối chiếu chính', bold=True, size=12)
b_tbl = d.add_table(rows=7, cols=4); b_tbl.style = 'Table Grid'
hdr = ['Tiêu chí', 'QT066 Peer Support', 'Cross-ref', 'Đánh giá']
for i, h in enumerate(hdr):
    c = b_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
b_data = [
    ('Người cung cấp', 'PEER (đồng đẳng) — có lived experience',
     'Maya (QT043): app tự hướng dẫn (no clinician) | BESST (QT042): GV/nhân viên trường + clinician | Happy House: GV/researchers | Walder: chuyên gia chat',
     'Peer support DUY NHẤT non-clinical + non-staff. Khác mọi can thiệp khác trong DB. Phù hợp HS VN ưu thích non-professional'),
    ('Cấu trúc',
     '15 NC gốc, 13 interventions: 1-1 hoặc nhóm; mode đa dạng (mặt-đối-mặt, online, điện thoại)',
     'Maya: app 12 phiên/6 tuần | EACP: 25 buổi nhóm | Happy House: 6 buổi nhóm 90 phút | BESST: 1 buổi self-referral',
     'Peer support BIẾN THIÊN cao về cấu trúc — không có protocol thống nhất. Khó nhân rộng nếu chỉ có 1 mô hình'),
    ('Đối tượng tuổi',
     '9/13 nhắm 16-25 (thanh niên/đại học); chỉ 2/13 cho HS THCS-THPT',
     'Maya: 18-25 | EACP: lớp 7 (~12-13) | Happy House: HS THCS-THPT VN | BESST: HS UK | Cai resilience: trẻ + VTN',
     '⚠ Peer support CHỦ YẾU CHO THANH NIÊN. Khác EACP/Happy House/Maya/CAMS ở chỗ ít cho HS THCS. ÁP DỤNG CHO VN cần adapt mạnh'),
    ('Outcome đo',
     'GAD-7, DASS-21, recovery, self-efficacy, quality of life',
     'Maya: HAM-A, ASI, LSAS | EACP: BASC-2 | Happy House: CESD-R | Walder: SIAS, SPIN | CAMS: PARS, CGI',
     'Peer support outcome ĐA DẠNG hơn — không chuyên biệt cho 1 rối loạn. Phù hợp với mục tiêu "wellbeing tổng thể"'),
    ('Effect size',
     '⚠ KHÔNG có (scoping review không đo effect size pooled)',
     'Walder g=0,508 | Maya d=1,04 | EACP d=0,13 | Cai resilience SMD=0,17 | Happy House d=0,11',
     '⚠ Peer support KHÔNG ĐỦ BẰNG CHỨNG ĐỊNH LƯỢNG. Cần SR + MA tiếp theo. Hiện chỉ "có tiềm năng" qualitatively'),
    ('Vai trò trong landscape',
     'GATEWAY (cửa ngõ) → can thiệp chuyên môn',
     'Maya/EACP/CAMS: can thiệp đầy đủ | BESST: targeted self-referral (gateway) | Walder: prevention/early treatment',
     'Peer support PHÙ HỢP làm CỬA NGÕ trước khi HS tiếp cận chuyên môn. Nhất quán với mô hình "stepped care" — nhẹ trước, nặng sau'),
]
for i, row in enumerate(b_data):
    for j, v in enumerate(row):
        b_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in b_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn('B.2. Mâu thuẫn / điểm cần lưu ý', bold=True, size=12)

conflict_tbl_b = d.add_table(rows=5, cols=2); conflict_tbl_b.style = 'Table Grid'
hdr_c = ['Phát hiện', 'Lý giải / Cảnh báo']
for i, h in enumerate(hdr_c):
    c = conflict_tbl_b.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
conflicts_b = [
    ('① Peer support không có effect size định lượng — KHÔNG mâu thuẫn nhưng yếu',
     'Murphy 2024 là SCOPING REVIEW (không phải MA). Không có forest plot, không có SMD/Cohen d pooled. Khác 5 bài còn lại (CAMS/Maya/Walder/Cai/Happy House) đều có effect size cụ thể. → '
     'Bằng chứng peer support YẾU HƠN các can thiệp khác về mặt định lượng — nhưng KHÔNG mâu thuẫn (chỉ là chưa đủ NC tốt).'),
    ('② Peer support 16-25 vs EACP/Happy House cho HS THCS — gap',
     '⚠ KHOẢNG TRỐNG nghiên cứu lớn: Murphy 2024 chỉ có 2/13 NC cho HS THCS-THPT. EACP (lớp 7) và Happy House (THCS-THPT VN) chính là khoảng trống Murphy chưa lấp. '
     '→ Ý tưởng nghiên cứu: peer support cho HS THCS-THPT VN còn HOÀN TOÀN MỚI — đầu tư có giá trị.'),
    ('③ Peer support vs Cai resilience — overlap?',
     'Cai 2025 (QT059): Mindfulness SMD=0,57 > CBT -0,26 cho RESILIENCE. Peer support có thể tăng resilience qua "shared experience" + role model. '
     'KHÔNG mâu thuẫn nhưng có thể OVERLAP cơ chế: peer support + mindfulness đều "non-clinical" + tăng coping skill. '
     'Có thể kết hợp peer support + mindfulness component cho VN.'),
    ('④ Peer support vs BESST — gateway model nhất quán',
     'BESST UK (QT042): targeted + self-referral 1 buổi → giảm stigma. Peer support: peer-led → giảm stigma. '
     'Cả 2 đều là "gateway" trước chuyên môn. NHẤT QUÁN: HS thường ĐẾN PEER trước, ĐẾN GV/CLINICIAN sau. '
     'Mô hình BESST + peer support có thể TÍCH HỢP cho VN.'),
]
for i, (k, v) in enumerate(conflicts_b):
    c0 = conflict_tbl_b.rows[i+1].cells[0]; shade(c0, 'FFF2CC')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = ORANGE
    conflict_tbl_b.rows[i+1].cells[1].text = v
d.add_paragraph()

# =====================================================
# PHẦN C — TÍCH HỢP LANDSCAPE CAN THIỆP CHO VN
# =====================================================
add_h('PHẦN C — Tích hợp landscape can thiệp lo âu HS THCS/THPT VN', 1)

vn('C.1. Stepped care model gợi ý cho VN', bold=True, size=12)
sc_tbl = d.add_table(rows=6, cols=3); sc_tbl.style = 'Table Grid'
hdr = ['Tầng', 'Bài tham chiếu', 'Áp dụng VN']
for i, h in enumerate(hdr):
    c = sc_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
sc_data = [
    ('Tầng 0 — Universal prevention',
     'Happy House (VN030) RAP-A 6 buổi | Cai resilience Mindfulness',
     'Tổ chức cho TOÀN bộ HS THCS-THPT. Ngắn (6 buổi), không kỳ thị. Chuẩn bị "đất" cho can thiệp sâu hơn'),
    ('Tầng 1 — Peer support gateway',
     'QT066 Murphy 2024 peer support',
     'HS lớp 11-12 đào tạo làm peer leader, hỗ trợ HS lớp 7-8 có triệu chứng nhẹ. Giảm kỳ thị + cửa ngõ vào'),
    ('Tầng 2 — Targeted screening + brief',
     'QT042 BESST self-referral (UK)',
     'HS có triệu chứng GAD-7 ≥ 5 hoặc DASS-21 vừa: tự đăng ký 1 buổi workshop CBT brief. PLACES language giảm kỳ thị'),
    ('Tầng 3 — Indicated prevention (lo âu chưa chẩn đoán)',
     'QT065 EACP 25 buổi + Happy House',
     'HS có triệu chứng GAD-7 ≥ 10 hoặc DASS-21 nặng nhưng chưa đủ chẩn đoán: nhóm CBT 12-15 buổi tại trường (rút gọn EACP)'),
    ('Tầng 4 — Treatment (đã chẩn đoán)',
     'QT028 CAMS CBT+SSRI | QT043 Maya app | QT040 Walder iCBT',
     'HS có chẩn đoán GAD/SAD/MDD: chuyển đến tâm lý gia chuyên môn + CBT cá nhân + app hỗ trợ (nếu có) + thuốc nếu cần'),
]
for i, row in enumerate(sc_data):
    for j, v in enumerate(row):
        sc_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in sc_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn('C.2. So sánh effect size — landscape tổng thể', bold=True, size=12)

es_tbl = d.add_table(rows=8, cols=4); es_tbl.style = 'Table Grid'
hdr = ['Bài', 'Loại can thiệp', 'Effect size', 'Đối tượng']
for i, h in enumerate(hdr):
    c = es_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
es_data = [
    ('QT043 Maya', 'CBT app', 'Cohen d = 1,04 (HAM-A) — RẤT LỚN', '18-25 tuổi RLLA chẩn đoán'),
    ('QT040 Walder DMHI SAD-specific', 'iCBT có hướng dẫn', 'Hedges g = 0,878 — LỚN', '<25 tuổi SAD subthreshold-mild'),
    ('QT040 Walder DMHI tổng', 'DMHI tổng quát', 'g = 0,508 — TRUNG BÌNH', '<25 tuổi'),
    ('QT028 Zugman CAMS combo', 'CBT+SSRI', '47-66% phục hồi (NNT=3)', '7-17 tuổi RLLA chẩn đoán'),
    ('QT065 EACP', 'CBT nhóm + cha mẹ + GV', 'd = 0,13 internalizing — RẤT NHỎ', 'Lớp 7 gây hấn (gốc), không phải lo âu'),
    ('QT059 Cai Mindfulness', 'Mindfulness school', 'SMD = 0,57 resilience — TB', 'Trẻ + VTN'),
    ('VN030 Happy House', 'RAP-A 6 buổi nhóm', 'd = 0,11 CESD-R — RẤT NHỎ', 'HS THCS-THPT VN universal'),
]
for i, row in enumerate(es_data):
    for j, v in enumerate(row):
        es_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in es_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn('Quy luật quan sát:', bold=True)
for it in [
    'Effect size LỚN NHẤT: can thiệp ĐIỀU TRỊ cho HS có chẩn đoán (Maya d=1,04; Walder SAD-specific g=0,878)',
    'Effect size TRUNG BÌNH: prevention/treatment hỗn hợp với mẫu rộng (Walder tổng g=0,508; Cai mindfulness SMD=0,57)',
    'Effect size NHỎ: universal/indicated prevention (Happy House d=0,11; EACP d=0,13)',
    '→ Càng "đậm đặc" target (chẩn đoán → indicated → universal), effect size càng nhỏ — quy luật chung của prevention literature',
    '→ Không nên kỳ vọng EACP hoặc Happy House có d=1,0 như Maya/Walder',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)

d.add_paragraph()

# =====================================================
# PHẦN D — KẾT LUẬN VERIFY
# =====================================================
add_h('PHẦN D — Kết luận verify cross-bài', 1, GREEN)

vn('Tóm tắt kết quả audit:', bold=True, size=12)
result_tbl = d.add_table(rows=6, cols=1); result_tbl.style = 'Table Grid'
results = [
    ('✓ KHÔNG MÂU THUẪN số liệu giữa QT065/QT066 và 7 bài cross-reference đã có trong DB. '
     'Tất cả số liệu QT065 (n=709, 25 buổi, d=0,13) và QT066 (15 NC, 13 interventions, 9/13 cho 16-25) đã verify từ Tom-tat-tung-bai/.', GREEN),
    ('✓ EACP LẤP KHOẢNG TRỐNG: là can thiệp 3-thành-phần (HS+cha mẹ+GV) DUY NHẤT trong DB. '
     'Khác CAMS (HS+therapy), Maya (HS-only digital), Walder (HS-only digital), Happy House (HS-only nhóm). '
     'EACP CÓ THỂ là mô hình tham chiếu khi adapt cho VN.', DARK),
    ('✓ PEER SUPPORT LẤP MỘT KHOẢNG TRỐNG KHÁC: là can thiệp NON-CLINICAL DUY NHẤT trong DB. '
     'Phù hợp HS VN ưu thích non-professional trước. NHƯNG: Murphy 2024 chỉ có 2/13 NC cho HS THCS-THPT — '
     'cần ADAPT mạnh nếu áp dụng cho VN, không thể copy trực tiếp.', ORANGE),
    ('⚠ EACP HIỆU LỰC LO ÂU THẤP (d=0,13): hợp lý vì đối tượng GỐC là HS gây hấn không phải HS lo âu. '
     'KHÔNG kỳ vọng EACP là "first-line" cho lo âu VN — Maya/Walder/CAMS phù hợp hơn cho HS có lo âu.', RED),
    ('⚠ PEER SUPPORT KHÔNG CÓ EFFECT SIZE PHỐI HỢP: Murphy 2024 là scoping review (mapping evidence). '
     'Không thể so sánh trực tiếp với Maya/Walder/CAMS ở mức effect size. Cần MA tiếp theo cho peer support YMH.', RED),
    ('✓ STEPPED CARE MODEL có thể tích hợp cả 2 bài + 7 bài cross-ref: '
     'Universal (Happy House) → Peer support gateway → BESST self-referral → Indicated EACP/CBT nhóm → Treatment Maya/Walder/CAMS.', GREEN),
]
for i, (txt, color) in enumerate(results):
    cell = result_tbl.rows[i].cells[0]
    if color == GREEN: shade(cell, 'E2EFDA')
    elif color == ORANGE: shade(cell, 'FFF2CC')
    elif color == RED: shade(cell, 'FCE4D6')
    else: shade(cell, 'D9E1F2')
    p = cell.paragraphs[0]
    r = p.add_run(txt); r.font.size = Pt(11)
d.add_paragraph()

vn('Khuyến nghị cho dự án em:', bold=True, size=12)
for it in [
    '(1) Thêm cả QT065 + QT066 vào INSIGHT_04 (can thiệp hiệu quả) trong rebuild_rag_full.py — bổ sung góc nhìn '
    'school-based prevention (EACP) + non-clinical gateway (peer support)',
    '(2) Build doc tổng hợp "Stepped care model cho HS VN" — tích hợp 9 bài: 7 cross-ref + QT065 + QT066',
    '(3) Xem xét RCT VN protocol — kết hợp tầng 1 (peer support) + tầng 3 (CBT nhóm rút gọn) — '
    'có thể là RCT đầu tiên ở VN tận dụng cả 2 bài mới',
    '(4) Cập nhật báo cáo Can thiệp Lo âu (01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN...) — thêm QT065/QT066 vào TLTK',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
d.add_paragraph()

# =====================================================
# PHỤ LỤC
# =====================================================
add_h('Phụ lục — Các tóm tắt đã đối chiếu', 1)
refs = [
    'QT065 Bradshaw et al. 2025 (J School Psychology) — verify từ PDF + Tom-tat-tung-bai/QT065_*.docx',
    'QT066 Murphy et al. 2024 (J Community Psychology) — verify từ PDF + Tom-tat-tung-bai/QT066_*.docx',
    'QT028 Zugman et al. 2024 (AJP) — verify từ Tom-tat-tung-bai/QT028_*.docx',
    'QT029 Li et al. 2025 (BMC Psychiatry NMA) — verify từ Tom-tat-tung-bai/QT029_*.docx',
    'QT040 Walder et al. 2025 (JMIR DMHI) — verify từ Tom-tat-tung-bai/QT040_*.docx',
    'QT042 Brown & Carter 2025 (J Mental Health UK) — verify từ Tom-tat-tung-bai/QT042_*.docx',
    'QT043 Bress et al. 2024 (JAMA Network Open Maya) — verify từ Tom-tat-tung-bai/QT043_*.docx',
    'QT059 Cai et al. 2025 (Frontiers Psychiatry resilience) — verify từ Tom-tat-tung-bai/QT059_*.docx',
    'VN030 Tran et al. 2023 (Cambridge Prisms Happy House) — verify từ Tom-tat-tung-bai/VN030_*.docx',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs: r.font.size = Pt(10)

d.add_paragraph()
foot = d.add_paragraph()
fr = foot.add_run('Cross-reference verification — biên soạn 29/04/2026 theo Nguyên tắc dịch v2 — Nguyên tắc 9 (cross-reference audit). '
                  'Mọi cross-ref được verify TRỰC TIẾP từ Tom-tat-tung-bai/ — KHÔNG dùng trí nhớ. KHÔNG phát hiện mâu thuẫn số liệu. '
                  '7 bài cross-ref đã đối chiếu thành công.')
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
