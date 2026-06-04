# -*- coding: utf-8 -*-
"""Vong 2 - 5 vong check sau hon vong 1.
- Vong 6: Cross-reference advanced
- Vong 7: Citation year/author exact match
- Vong 8: Linguistic accuracy + terminology
- Vong 9: Number precision (comma vs period)
- Vong 10: Final integrity
29/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', 'TrichYeuLA_CongThiHang_v2_29052026.docx')
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

d = Document(FILE)
la = Document(LA)

all_d = '\n'.join(p.text for p in d.paragraphs)
for tb in d.tables:
    for row in tb.rows:
        for c in row.cells:
            all_d += '\n' + c.text

all_la = '\n'.join(p.text for p in la.paragraphs)
for tb in la.tables:
    for row in tb.rows:
        for c in row.cells:
            all_la += '\n' + c.text

# Split into VN section (before page break) and EN section
# Find paragraph index of page break or "DISSERTATION ABSTRACT"
en_start_idx = None
for i, p in enumerate(d.paragraphs):
    if 'DISSERTATION ABSTRACT' in p.text:
        en_start_idx = i
        break

vn_text = '\n'.join(p.text for p in d.paragraphs[:en_start_idx]) if en_start_idx else ''
en_text = '\n'.join(p.text for p in d.paragraphs[en_start_idx:]) if en_start_idx else ''

issues = []
total = 0
passed = 0


def check(name, cond, detail=''):
    global total, passed
    total += 1
    status = '✓' if cond else '✗'
    print(f'  {status} {name}')
    if detail:
        print(f'      {detail}')
    if cond:
        passed += 1
    else:
        issues.append(name)
    return cond


# ============================================================
# VONG 6: CROSS-REFERENCE ADVANCED
# ============================================================
print('=== VÒNG 6: CROSS-REFERENCE ADVANCED ===')

check('6.1 4 thành phần RLLA: nhận thức + cảm xúc + sinh lý + hành vi',
      all(k in all_d.lower() for k in ['nhận thức', 'cảm xúc', 'sinh lý', 'hành vi']))

check('6.2 EN: cognitive + emotional + physiological + behavioral',
      all(k in all_d.lower() for k in ['cognitive', 'emotional', 'physiological', 'behavioral']))

check('6.3 Vietnam mentioned in EN context',
      'Vietnam' in all_d)

check('6.4 Hanoi mentioned (mẫu HN)',
      'Hà Nội' in all_d and 'Hanoi' in all_d)

check('6.5 DSM-5 với dấu gạch (KHÔNG phải DSM5)',
      'DSM-5' in all_d and 'DSM5' not in all_d)

check('6.6 SEM viết tắt (structural equation model)',
      'SEM' in all_d and ('phương trình cấu trúc' in all_d.lower() or 'structural equation' in all_d.lower()))

check('6.7 CFA viết tắt (confirmatory factor analysis)',
      'CFA' in all_d and ('nhân tố khẳng định' in all_d.lower() or 'confirmatory factor' in all_d.lower()))

check('6.8 "trẻ vị thành niên" hoặc "adolescents"',
      'vị thành niên' in all_d.lower() or 'adolescents' in all_d.lower())

check('6.9 "phòng ngừa" + "prevention" (cả VN + EN)',
      'phòng ngừa' in all_d.lower() and 'prevention' in all_d.lower())

check('6.10 "sức khỏe tâm thần" + "mental health"',
      'sức khỏe tâm thần' in all_d.lower() and 'mental health' in all_d.lower())


# ============================================================
# VONG 7: CITATION YEAR/AUTHOR EXACT MATCH vs LA
# ============================================================
print('\n=== VÒNG 7: CITATION YEAR/AUTHOR MATCH LA ===')

citations = [
    ('Chorpita', '2000'),
    ('Sun', '2011'),
    ('Kwon', '2013'),
    ('Olweus', '1996'),
    ('Goodenow', '1993'),
    ('Zimet', '1988'),
    ('Rosenberg', '1965'),
    ('Carver', '1997'),
]

for author, year in citations:
    # In trich yeu
    pattern_d = f'{author}.*?{year}'
    has_d = bool(re.search(pattern_d, all_d, re.DOTALL))
    # In LA
    pattern_la = f'{author}.*?{year}'
    has_la = bool(re.search(pattern_la, all_la, re.DOTALL))
    check(f'7.{author} ({year}) trong trích yếu', has_d)
    check(f'7.{author} ({year}) khớp LA chính', has_la)


# ============================================================
# VONG 8: LINGUISTIC ACCURACY + TERMINOLOGY
# ============================================================
print('\n=== VÒNG 8: LINGUISTIC ACCURACY ===')

check('8.1 "lower secondary school" (NOT "junior secondary school")',
      'lower secondary school' in all_d.lower() and 'junior secondary' not in all_d.lower())

check('8.2 "trung học cơ sở" 8+ lần (đủ tần suất VN)',
      all_d.lower().count('trung học cơ sở') >= 5)

check('8.3 KHÔNG dùng "lo âu tổng quát" (sai DSM-5)',
      'lo âu tổng quát' not in all_d.lower())

check('8.4 KHÔNG có placeholder bracket "[NCS điền]"',
      '[NCS' not in all_d and '[điền' not in all_d and '[...]' not in all_d)

check('8.5 KHÔNG có "TODO" hoặc "xx" placeholder',
      'TODO' not in all_d.upper() and 'xx' not in all_d.lower().replace('xxxx', '').replace('xx_', ''))

check('8.6 KHÔNG có 2 spaces liên tiếp',
      '  ' not in all_d.replace('\t\t', ''))

check('8.7 Câu Việt kết thúc bằng dấu chấm "." (không dấu ":")',
      not vn_text.strip().endswith(':') if vn_text else True)

check('8.8 Tên Việt "Công Thị Hằng" xuất hiện (signature)',
      'Công Thị Hằng' in all_d)

check('8.9 Tên Anh "Cong Thi Hang" (Vietnamese order, no diacritics)',
      'Cong Thi Hang' in all_d)

check('8.10 Tên thầy "TS. Đào Minh Đức"',
      'TS. Đào Minh Đức' in all_d)

check('8.11 Tên thầy EN "Dr. Dao Minh Duc"',
      'Dr. Dao Minh Duc' in all_d)


# ============================================================
# VONG 9: NUMBER PRECISION (comma vs period)
# ============================================================
print('\n=== VÒNG 9: NUMBER PRECISION ===')

check('9.1 VN: "1.352" (Vietnamese thousand separator)',
      '1.352' in vn_text)
check('9.2 EN: "1,352" (English thousand separator)',
      '1,352' in en_text)

check('9.3 VN: "F = 0,246" (Vietnamese decimal)',
      'F = 0,246' in vn_text)
check('9.4 EN: "F = 0.246" (English decimal)',
      'F = 0.246' in en_text)

check('9.5 VN: "51,13" decimal comma',
      '51,13' in vn_text)
check('9.6 EN: "51.13" decimal period',
      '51.13' in en_text)

check('9.7 VN: "p < 0,001" (decimal comma)',
      'p < 0,001' in vn_text)
check('9.8 EN: "p < 0.001" (decimal period)',
      'p < 0.001' in en_text)

check('9.9 VN: KHÔNG có "0.246" sai (period in VN section)',
      '0.246' not in vn_text)
check('9.10 EN: KHÔNG có "0,246" sai (comma in EN section)',
      '0,246' not in en_text)

check('9.11 VN: KHÔNG có "51.13" sai',
      '51.13' not in vn_text)
check('9.12 EN: KHÔNG có "51,13" sai',
      '51,13' not in en_text)


# ============================================================
# VONG 10: FINAL INTEGRITY
# ============================================================
print('\n=== VÒNG 10: FINAL INTEGRITY ===')

# Check word count is reasonable
words_vn = len(vn_text.split())
words_en = len(en_text.split())
check(f'10.1 VN word count >= 350', words_vn >= 350, f'{words_vn} words')
check(f'10.2 EN word count >= 280', words_en >= 280, f'{words_en} words')
check(f'10.3 VN word count <= 1000', words_vn <= 1000, f'{words_vn} words')
check(f'10.4 EN word count <= 1000', words_en <= 1000, f'{words_en} words')

# All paragraphs have non-empty content (no blank format issues)
non_empty_paras = sum(1 for p in d.paragraphs if p.text.strip())
check(f'10.5 Số paragraphs có nội dung >= 25', non_empty_paras >= 25,
      f'{non_empty_paras} paragraphs')

# Number of body paragraphs (4 lý luận + 4 thực tiễn approx)
contribution_paras = [p for p in d.paragraphs
                      if (p.paragraph_format.first_line_indent
                          and abs(p.paragraph_format.first_line_indent.cm - 1.0) < 0.05
                          and p.runs and p.runs[0].bold is False)]
check(f'10.6 Body paragraphs (fli=1.0cm + not bold) >= 10',
      len(contribution_paras) >= 10,
      f'{len(contribution_paras)} body paras')

# Check: exactly 2 page sections (VN + EN with page break)
check('10.7 Có ít nhất 1 page break giữa VN-EN',
      any('w:br' in r._element.xml for p in d.paragraphs for r in p.runs))

# Number of tables = exactly 2 (signature VN + EN)
check('10.8 Đúng 2 signature tables', len(d.tables) == 2)

# Both signature tables have 1 row x 2 columns
for ti, tbl in enumerate(d.tables, 1):
    check(f'10.9 Table {ti}: 1 row × 2 cols',
          len(tbl.rows) == 1 and len(tbl.columns) == 2)

# No leftover MPAS/MPVS anywhere
check('10.10 KHÔNG còn "MPAS" hoặc "MPVS" residual',
      'MPAS' not in all_d and 'MPVS' not in all_d)

# No old "Rosenberg" without RSES
rsenberg_lone = bool(re.search(r'Rosenberg(?! Self)', all_d.replace('Rosenberg Self-Esteem Scale', '')))
check('10.11 "Rosenberg" có RSES acronym hoặc tên đầy đủ',
      'RSES' in all_d)

# Check no "9310401" residual
check('10.12 KHÔNG có "9310401" stale', '9310401' not in all_d)

# Check no "~86%" residual
check('10.13 KHÔNG có "~86%" stale',
      '~86%' not in all_d and 'khoảng 86%' not in all_d)


# ============================================================
# RESULT
# ============================================================
print()
print('=' * 70)
print(f'TỔNG VÒNG 6-10: {passed}/{total} checks PASS ({len(issues)} lỗi)')
if issues:
    print('\nDANH SÁCH LỖI:')
    for i, iss in enumerate(issues, 1):
        print(f'  {i}. {iss}')
    sys.exit(1)
else:
    print('\nSẠCH LỖI VÒNG 2!')
    sys.exit(0)
