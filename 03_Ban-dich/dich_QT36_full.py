# -*- coding: utf-8 -*-
"""Dich day du QT36 - Korea GAD ML 2025 - Moon & Woo - Frontiers in Public Health"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.3389/fpubh.2024.1504739', size=10)
add_p(doc, 'PMC: https://pmc.ncbi.nlm.nih.gov/articles/PMC11746110/', size=10)

add_heading(doc, 'C\u00e1c y\u1ebfu t\u1ed1 nguy c\u01a1 ch\u00ednh c\u1ee7a r\u1ed1i lo\u1ea1n lo \u00e2u t\u1ed5ng qu\u00e1t \u1edf thanh thi\u1ebfu ni\u00ean: Nghi\u00ean c\u1ee9u h\u1ecdc m\u00e1y', 1)
h = doc.add_paragraph()
r = h.add_run('Key risk factors of generalized anxiety disorder in adolescents: machine learning study')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('Ti\u00eau \u0111\u1ec1 g\u1ed1c', 'Key risk factors of generalized anxiety disorder in adolescents: machine learning study'),
    ('Ti\u00eau \u0111\u1ec1 d\u1ecbch', 'C\u00e1c y\u1ebfu t\u1ed1 nguy c\u01a1 ch\u00ednh c\u1ee7a r\u1ed1i lo\u1ea1n lo \u00e2u t\u1ed5ng qu\u00e1t \u1edf thanh thi\u1ebfu ni\u00ean: Nghi\u00ean c\u1ee9u h\u1ecdc m\u00e1y'),
    ('T\u00e1c gi\u1ea3', 'Yonghwan Moon (1), Hyekyung Woo (1,2)*'),
    ('C\u01a1 quan', '(1) Khoa Qu\u1ea3n l\u00fd Y t\u1ebf, \u0110H Qu\u1ed1c gia Kongju, H\u00e0n Qu\u1ed1c\n(2) Vi\u1ec7n Y t\u1ebf v\u00e0 M\u00f4i tr\u01b0\u1eddng, \u0110H Qu\u1ed1c gia Kongju'),
    ('T\u1ea1p ch\u00ed', 'Frontiers in Public Health (Q1, IF \u2248 5,2)'),
    ('Xu\u1ea5t b\u1ea3n', '2025, 9 trang'),
    ('DOI', '10.3389/fpubh.2024.1504739'),
    ('Lo\u1ea1i NC', 'Ph\u00e2n t\u00edch d\u1eef li\u1ec7u th\u1ee9 c\u1ea5p \u2014 Machine Learning (LASSO, SelectKBest, XGBoost, Random Forest, ANN)'),
    ('M\u1eabu', '213.820 VTN H\u00e0n Qu\u1ed1c (KYRBS 2020\u20132023); THCS 53,8% + THPT 46,2%'),
    ('Truy c\u1eadp', 'Open Access \u2014 Frontiers'),
])
add_page_ref(doc, '1\u20139', 'Frontiers in Public Health', '2025')

add_heading(doc, 'T\u00d3M T\u1eaeT', 2)
add_p(doc, 'Thanh thi\u1ebfu ni\u00ean tr\u00ean to\u00e0n th\u1ebf gi\u1edbi ng\u00e0y c\u00e0ng b\u1ecb \u1ea3nh h\u01b0\u1edfng b\u1edfi c\u00e1c r\u1ed1i lo\u1ea1n s\u1ee9c kh\u1ecfe t\u00e2m th\u1ea7n, trong \u0111\u00f3 R\u1ed1i lo\u1ea1n Lo \u00e2u T\u1ed5ng qu\u00e1t (GAD \u2014 Generalized Anxiety Disorder) \u0111\u1eb7c bi\u1ec7t ph\u1ed5 bi\u1ebfn. M\u1eb7c d\u00f9 t\u00e1c \u0111\u1ed9ng \u0111\u00e1ng k\u1ec3, GAD \u1edf thanh thi\u1ebfu ni\u00ean th\u01b0\u1eddng b\u1ecb ch\u1ea9n \u0111o\u00e1n thi\u1ebfu do tri\u1ec7u ch\u1ee9ng m\u01a1 h\u1ed3 v\u00e0 ch\u1eadm tr\u1ec5 trong vi\u1ec7c t\u00ecm ki\u1ebfm h\u1ed7 tr\u1ee3 y t\u1ebf.')

p = doc.add_paragraph()
r = p.add_run('Ph\u01b0\u01a1ng ph\u00e1p: S\u1eed d\u1ee5ng d\u1eef li\u1ec7u KYRBS (Korea Youth Risk Behavior Web-based Survey) 2020\u20132023 \u0111\u1ec3 ph\u00e2n t\u00edch c\u00e1c y\u1ebfu t\u1ed1 \u1ea3nh h\u01b0\u1edfng \u0111\u1ebfn GAD \u1edf VTN. K\u1ef9 thu\u1eadt h\u1ecdc m\u00e1y: H\u1ed3i quy Lasso, SelectKBest, v\u00e0 XGBoost \u0111\u1ec3 ch\u1ecdn \u0111\u1eb7c tr\u01b0ng. M\u00f4 h\u00ecnh d\u1ef1 \u0111o\u00e1n: Random Forest v\u00e0 M\u1ea1ng N\u01a1-ron Nh\u00e2n t\u1ea1o (ANN). SMOTE x\u1eed l\u00fd m\u1ea5t c\u00e2n b\u1eb1ng. n = 213.820 VTN (THCS 53,8% + THPT 46,2%). GAD \u0111o b\u1eb1ng GAD-7 \u2265 10.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('K\u1ebft qu\u1ea3: GAD prevalence = 17,9%. Top y\u1ebfu t\u1ed1: (1) C\u00f4 \u0111\u01a1n \u2014 LASSO 0,23, F-score 16.646, XGBoost gain 0,29 (M\u1ea0NH NH\u1ea4T c\u1ea3 3 thu\u1eadt to\u00e1n); (2) Stress \u2014 0,10, 6.589; (3) Nh\u1eadn th\u1ee9c s\u1ee9c kh\u1ecfe ch\u1ee7 quan k\u00e9m; (4) Ph\u1ee5c h\u1ed3i gi\u1ea5c ng\u1ee7 k\u00e9m; (5) B\u1ea1o l\u1ef1c. RF (XGBoost features): accuracy 78%, AUC 82%, sensitivity 64%. ANN: accuracy 80%, AUC 81%.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RGBColor(0xCC, 0, 0)

add_p(doc, 'K\u1ebft lu\u1eadn: C\u00e1c ph\u00e1t hi\u1ec7n nh\u1ea5n m\u1ea1nh nhu c\u1ea7u ch\u01b0\u01a1ng tr\u00ecnh gi\u00e1o d\u1ee5c t\u1eadp trung v\u00e0o qu\u1ea3n l\u00fd gi\u1ea5c ng\u1ee7, ph\u00f2ng ng\u1eeba h\u00fat thu\u1ed1c, v\u00e0 dinh d\u01b0\u1ee1ng c\u00e2n b\u1eb1ng \u0111\u1ec3 gi\u1ea3m nguy c\u01a1 GAD \u1edf VTN.')
add_p(doc, 'T\u1eeb kh\u00f3a: thanh thi\u1ebfu ni\u00ean, s\u1ee9c kh\u1ecfe t\u00e2m th\u1ea7n, r\u1ed1i lo\u1ea1n lo \u00e2u t\u1ed5ng qu\u00e1t, h\u1ecdc m\u00e1y, h\u00e0nh vi s\u1ee9c kh\u1ecfe.')

# DANH GIA NHANH
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'Frontiers in Public Health Q1 IF \u2248 5,2. Machine Learning \u2014 ph\u01b0\u01a1ng ph\u00e1p M\u1edaI trong NC GAD VTN.',
    'M\u1eabu C\u1ef0C L\u1edaN: 213.820 VTN H\u00e0n Qu\u1ed1c (KYRBS 2020\u20132023).',
    'GAD-7 \u0111o lo \u00e2u T\u1ed4NG QU\u00c1T ri\u00eang (kh\u00f4ng g\u1ed9p chung DAS) \u2014 ch\u00ednh x\u00e1c h\u01a1n DASS-21.',
    'C\u00d4 \u0110\u01a0N x\u1ebfp h\u1ea1ng 1 \u2014 ph\u00e1t hi\u1ec7n quan tr\u1ecdng, v\u01b0\u1ee3t tr\u1ed9i stress v\u00e0 gi\u1ea5c ng\u1ee7.',
    'H\u00c0NH VI S\u1ee8C KH\u1ece (gi\u1ea5c ng\u1ee7, h\u00fat thu\u1ed1c, th\u1ee9c \u0103n nhanh) \u2014 y\u1ebfu t\u1ed1 C\u00d3 TH\u1ec2 CAN THI\u1ec6P.',
    'THCS 53,8% + THPT 46,2% \u2014 bao ph\u1ee7 C\u1ea2 c\u1ea5p 2 v\u00e0 c\u1ea5p 3.',
    '3 thu\u1eadt to\u00e1n ch\u1ecdn \u0111\u1eb7c tr\u01b0ng + 2 m\u00f4 h\u00ecnh = 6 t\u1ed5 h\u1ee3p \u2014 so s\u00e1nh \u0111a ph\u01b0\u01a1ng ph\u00e1p.',
]:
    add_p(doc, '\u2022 ' + b)
add_p(doc, 'H\u1ea1n ch\u1ebf:', bold=True)
for b in [
    'Ch\u1ec9 H\u00e0n Qu\u1ed1c \u2014 kh\u00e1c VN v\u1ec1 gi\u00e1o d\u1ee5c, kinh t\u1ebf, v\u0103n h\u00f3a.',
    'C\u1eaft ngang \u2014 ML x\u00e1c \u0111\u1ecbnh y\u1ebfu t\u1ed1 li\u00ean quan, KH\u00d4NG ph\u1ea3i nh\u00e2n qu\u1ea3.',
    'SMOTE c\u00f3 th\u1ec3 g\u00e2y overfitting. Sensitivity ch\u1ec9 64% (36% b\u1ecf s\u00f3t).',
    'Kh\u00f4ng \u0111o: \u00e1p l\u1ef1c h\u1ecdc t\u1eadp c\u1ee5 th\u1ec3 (ESSA), screen time/MXH, ACEs.',
]:
    add_p(doc, '\u2022 ' + b)

# 1. GIOI THIEU
add_page_ref(doc, '1\u20132', 'Frontiers in Public Health', '2025')
add_heading(doc, '1. GI\u1edaI THI\u1ec6U', 2)
add_p(doc, 'Theo WHO, 14% VTN tr\u00ean to\u00e0n th\u1ebf gi\u1edbi tr\u1ea3i nghi\u1ec7m r\u1ed1i lo\u1ea1n t\u00e2m th\u1ea7n. Qu\u1ea3n l\u00fd kh\u00f4ng \u0111\u1ea7y \u0111\u1ee7 c\u00f3 th\u1ec3 d\u1eabn \u0111\u1ebfn gi\u00e1n \u0111o\u1ea1n nghi\u00eam tr\u1ecdng ch\u1ee9c n\u0103ng x\u00e3 h\u1ed9i, v\u1edbi h\u1eadu qu\u1ea3 nh\u01b0 t\u1ef1 t\u1eed. Trong c\u00e1c r\u1ed1i lo\u1ea1n t\u00e2m th\u1ea7n, r\u1ed1i lo\u1ea1n lo \u00e2u \u0111\u1eb7c bi\u1ec7t ph\u1ed5 bi\u1ebfn: 3,6% \u1edf nh\u00f3m 10\u201314 tu\u1ed5i v\u00e0 4,6% \u1edf nh\u00f3m 15\u201319 tu\u1ed5i \u2014 t\u1ef7 l\u1ec7 cao h\u01a1n ADHD v\u00e0 tr\u1ea7m c\u1ea3m.')
add_p(doc, 'R\u1ed1i lo\u1ea1n Lo \u00e2u T\u1ed5ng qu\u00e1t (GAD) th\u01b0\u1eddng xu\u1ea5t hi\u1ec7n trong tu\u1ed5i VTN v\u00e0 \u0111\u01b0\u1ee3c \u0111\u1eb7c tr\u01b0ng b\u1edfi lo l\u1eafng v\u00e0 lo \u00e2u dai d\u1eb3ng trong cu\u1ed9c s\u1ed1ng h\u00e0ng ng\u00e0y. Theo DSM-5, GAD \u0111\u01b0\u1ee3c ch\u1ea9n \u0111o\u00e1n khi lo \u00e2u qu\u00e1 m\u1ee9c k\u00e9o d\u00e0i \u00edt nh\u1ea5t 6 th\u00e1ng, k\u00e8m tri\u1ec7u ch\u1ee9ng: b\u1ed3n ch\u1ed3n, d\u1ec5 m\u1ec7t, kh\u00f3 t\u1eadp trung, c\u00e1u g\u1eaft, c\u0103ng c\u01a1, r\u1ed1i lo\u1ea1n gi\u1ea5c ng\u1ee7. N\u1ebfu kh\u00f4ng \u0111\u01b0\u1ee3c qu\u1ea3n l\u00fd, GAD c\u00f3 th\u1ec3 c\u1ea3n tr\u1edf t\u0103ng tr\u01b0\u1edfng v\u00e0 ph\u00e1t tri\u1ec3n su\u1ed1t \u0111\u1eddi.')
add_p(doc, 'Nhi\u1ec1u y\u1ebfu t\u1ed1 \u1ea3nh h\u01b0\u1edfng GAD \u1edf VTN \u0111\u00e3 \u0111\u01b0\u1ee3c b\u00e1o c\u00e1o: gi\u1edbi t\u00ednh, t\u00ecnh tr\u1ea1ng kinh t\u1ebf x\u00e3 h\u1ed9i, stress, nh\u1eadn th\u1ee9c s\u1ee9c kh\u1ecfe ch\u1ee7 quan. G\u1ea7n \u0111\u00e2y, h\u1ecdc m\u00e1y (ML) n\u1ed5i l\u00ean nh\u01b0 c\u00f4ng c\u1ee5 bi\u1ebfn \u0111\u1ed5i trong NC SKTT. ML \u0111\u00e3 ph\u00e2n bi\u1ec7t GAD v\u00e0 SAD v\u1edbi \u0111\u1ed9 ch\u00ednh x\u00e1c l\u00ean \u0111\u1ebfn 94% b\u1eb1ng fMRI. Tuy nhi\u00ean, \u1ee9ng d\u1ee5ng ML trong hi\u1ec3u y\u1ebfu t\u1ed1 nguy c\u01a1 GAD ri\u00eang \u1edf VTN v\u1eabn ch\u01b0a \u0111\u01b0\u1ee3c kh\u00e1m ph\u00e1 \u0111\u1ea7y \u0111\u1ee7.')
add_p(doc, 'Nghi\u00ean c\u1ee9u n\u00e0y \u0111i\u1ec1u tra c\u00e1c y\u1ebfu t\u1ed1 li\u00ean quan \u0111\u1ebfn GAD \u1edf VTN s\u1eed d\u1ee5ng d\u1eef li\u1ec7u KYRBS t\u1eeb 2020 \u0111\u1ebfn 2023, \u00e1p d\u1ee5ng ML \u0111\u1ec3 x\u00e1c \u0111\u1ecbnh c\u00e1c y\u1ebfu t\u1ed1 ch\u00ednh v\u00e0 ph\u00e1t tri\u1ec3n m\u00f4 h\u00ecnh d\u1ef1 \u0111o\u00e1n.')

# 2. PHUONG PHAP
add_page_ref(doc, '2\u20133', 'Frontiers in Public Health', '2025')
add_heading(doc, '2. PH\u01af\u01a0NG PH\u00c1P', 2)
add_p(doc, '2.1. D\u1eef li\u1ec7u', bold=True)
add_p(doc, 'KYRBS 2020\u20132023, qu\u1ea3n l\u00fd b\u1edfi KDCA (Korea Disease Control and Prevention Agency) v\u00e0 B\u1ed9 Gi\u00e1o d\u1ee5c H\u00e0n Qu\u1ed1c. Kh\u1ea3o s\u00e1t tr\u1ef1c tuy\u1ebfn t\u1ef1 b\u00e1o c\u00e1o, \u0111\u1ea1i di\u1ec7n qu\u1ed1c gia. N = 214.344 \u2192 lo\u1ea1i 524 (b\u00e1o c\u00e1o 0 gi\u1edd ng\u1ee7) \u2192 n = 213.820. D\u1eef li\u1ec7u c\u00f4ng khai t\u1ea1i https://www.kdca.go.kr/yhs/.')
add_p(doc, '2.2. Bi\u1ebfn ph\u1ee5 thu\u1ed9c', bold=True)
add_p(doc, 'GAD \u0111o b\u1eb1ng GAD-7 (Spitzer et al., 2006). Nh\u00f3m nguy c\u01a1 cao: GAD-7 \u2265 10 (moderate\u2013severe). T\u1ef7 l\u1ec7: 17,9% (38.288/213.820). Nh\u00f3m nguy c\u01a1 th\u1ea5p: 82,1% (175.532).')
add_p(doc, '2.3. Bi\u1ebfn \u0111\u1ed9c l\u1eadp', bold=True)
add_p(doc, 'Nh\u00e2n kh\u1ea9u: gi\u1edbi t\u00ednh, kh\u1ed1i (THCS/THPT), kinh t\u1ebf ch\u1ee7 quan, h\u1ecdc l\u1ef1c ch\u1ee7 quan, t\u00ecnh tr\u1ea1ng s\u1ed1ng, h\u1ecdc v\u1ea5n cha m\u1eb9. S\u1ee9c kh\u1ecfe: s\u1ee9c kh\u1ecfe ch\u1ee7 quan, ph\u1ee5c h\u1ed3i gi\u1ea5c ng\u1ee7, th\u1ee9c \u0103n nhanh, n\u01b0\u1edbc, h\u00fat thu\u1ed1c, r\u01b0\u1ee3u, b\u1ea1o l\u1ef1c, vi\u00eam m\u0169i d\u1ecb \u1ee9ng, stress, c\u00f4 \u0111\u01a1n, kinh nghi\u1ec7m t\u00ecnh d\u1ee5c, gi\u1ea5c ng\u1ee7 cu\u1ed1i tu\u1ea7n.')
add_p(doc, '2.4. SMOTE', bold=True)
add_p(doc, 'SMOTE \u0111\u01b0\u1ee3c \u00e1p d\u1ee5ng \u0111\u1ec3 x\u1eed l\u00fd m\u1ea5t c\u00e2n b\u1eb1ng l\u1edbp (17,9% GAD vs 82,1% kh\u00f4ng GAD). SMOTE t\u1ea1o m\u1eabu t\u1ed5ng h\u1ee3p cho l\u1edbp thi\u1ec3u s\u1ed1. \u0110\u1ec3 gi\u1ea3m overfitting, s\u1eed d\u1ee5ng ch\u1ec9 s\u1ed1 sensitivity, precision, v\u00e0 c\u00e1c ph\u01b0\u01a1ng ph\u00e1p regularization (LASSO) + ensemble (RF).')
add_p(doc, '2.5. Ch\u1ecdn \u0111\u1eb7c tr\u01b0ng', bold=True)
add_p(doc, '(1) LASSO \u2014 gi\u1ea3m \u0111a c\u1ed9ng tuy\u1ebfn, h\u1ec7 s\u1ed1 0 cho bi\u1ebfn kh\u00f4ng li\u00ean quan. (2) SelectKBest \u2014 \u01b0u ti\u00ean bi\u1ebfn t\u01b0\u01a1ng quan m\u1ea1nh nh\u1ea5t. (3) XGBoost \u2014 ph\u00e1t hi\u1ec7n quan h\u1ec7 phi tuy\u1ebfn v\u00e0 t\u01b0\u01a1ng t\u00e1c gi\u1eefa bi\u1ebfn.')
add_p(doc, '2.6. M\u00f4 h\u00ecnh d\u1ef1 \u0111o\u00e1n', bold=True)
add_p(doc, 'Random Forest (RF) v\u00e0 M\u1ea1ng N\u01a1-ron Nh\u00e2n t\u1ea1o (ANN). T\u1ed1i \u01b0u si\u00eau tham s\u1ed1 b\u1eb1ng grid search. 5-fold cross-validation. RF: max_depth=None, max_features=sqrt, min_samples_leaf=2, n_estimators=200. ANN: 2 hidden layers (64, 32), activation=relu, optimizer=adam, epochs=50.')

# 3. KET QUA
add_page_ref(doc, '3\u20136', 'Frontiers in Public Health', '2025')
add_heading(doc, '3. K\u1ebeT QU\u1ea2', 2)

add_p(doc, '3.1. \u0110\u1eb7c \u0111i\u1ec3m m\u1eabu (n = 213.820)', bold=True)
add_p(doc, 'Nam 61,8% (132.211), n\u1eef 38,2% (81.609). THCS 53,8% (114.966), THPT 46,2% (98.854). Kinh t\u1ebf ch\u1ee7 quan: cao 12,0%, trung b\u00ecnh 47,4%, th\u1ea5p 40,7%. H\u1ecdc l\u1ef1c ch\u1ee7 quan: cao 37,5%, trung b\u00ecnh 30,0%, th\u1ea5p 32,2%. S\u1ed1ng c\u00f9ng gia \u0111\u00ecnh 95,2%. H\u1ecdc v\u1ea5n cha: \u0110H+ 44,4%, THPT 16,2%, d\u01b0\u1edbi THCS 1,0%. S\u1ee9c kh\u1ecfe ch\u1ee7 quan: cao 65,7%, trung b\u00ecnh 25,1%, th\u1ea5p 9,3%. Ph\u1ee5c h\u1ed3i gi\u1ea5c ng\u1ee7 kh\u00f4ng \u0111\u1ea7y \u0111\u1ee7 41,1%. Th\u1ee9c \u0103n nhanh 1\u20132 l\u1ea7n/tu\u1ea7n 56,9%. H\u00fat thu\u1ed1c 9,3%. B\u1ea1o l\u1ef1c 9,4%. Vi\u00eam m\u0169i d\u1ecb \u1ee9ng 41,6%. Stress cao 81,2%. C\u00f4 \u0111\u01a1n 52,6%. Kinh nghi\u1ec7m t\u00ecnh d\u1ee5c 27,0%. GAD prevalence 17,9%.')

add_heading(doc, 'B\u1ea3ng 1. Top 11 y\u1ebfu t\u1ed1 nguy c\u01a1 GAD \u2014 Machine Learning', 3)
add_table(doc,
    ['H\u1ea1ng', 'Y\u1ebfu t\u1ed1', 'LASSO Coef.', 'SelectKBest F', 'XGBoost Gain', '\u00dd ngh\u0129a'],
    [['1', 'C\u00f4 \u0111\u01a1n', '0,23', '16.646', '0,29', 'M\u1ea0NH NH\u1ea4T \u2014 c\u1ea3 3 thu\u1eadt to\u00e1n'],
     ['2', 'Stress', '0,10', '6.589', '0,07', 'Nh\u1ea5t qu\u00e1n'],
     ['3', 'Nh\u1eadn th\u1ee9c s\u1ee9c kh\u1ecfe k\u00e9m', '0,12', '6.782', '0,04', ''],
     ['4', 'Ph\u1ee5c h\u1ed3i gi\u1ea5c ng\u1ee7 k\u00e9m', '0,10', '5.081', '0,02', 'Ph\u00f9 h\u1ee3p Zhu 2025 (QT05)'],
     ['5', 'B\u1ea1o l\u1ef1c', '0,10', '5.117', '0,01', ''],
     ['6', 'Kinh t\u1ebf th\u1ea5p', '0,03', '1.381', '0,01', 'Ph\u00f9 h\u1ee3p Korea 2024 (QT34)'],
     ['7', 'Ng\u1ee7 cu\u1ed1i tu\u1ea7n k\u00e9m', '0,03', '1.003', '0,01', ''],
     ['8', 'H\u00fat thu\u1ed1c', '0,03', '643', '0,01', ''],
     ['9', 'Kinh nghi\u1ec7m t\u00ecnh d\u1ee5c', '0,20', '11.219', '\u2014', 'LASSO+SelectKBest cao; XGBoost kh\u00f4ng ch\u1ecdn'],
     ['10', '\u0102n th\u1ee9c \u0103n nhanh h\u00e0ng ng\u00e0y', '\u2014', '\u2014', '0,02', 'Ch\u1ec9 XGBoost ch\u1ecdn'],
     ['11', 'N\u1eef gi\u1edbi', '\u2014', '\u2014', '0,004', 'Ch\u1ec9 XGBoost ch\u1ecdn']],
    widths=[0.8, 3.5, 2.0, 2.5, 2.0, 3.5])
add_p(doc, 'C\u00d4 \u0110\u01a0N v\u01b0\u1ee3t tr\u1ed9i stress \u1edf c\u1ea3 3 thu\u1eadt to\u00e1n. F-score = 16.646 (g\u1ea5p 2,5 l\u1ea7n stress). Gi\u1ea5c ng\u1ee7 (h\u1ea1ng 4+7) + h\u00fat thu\u1ed1c (h\u1ea1ng 8) + th\u1ee9c \u0103n nhanh (h\u1ea1ng 10) = H\u00c0NH VI S\u1ee8C KH\u1ece C\u00d3 TH\u1ec2 CAN THI\u1ec6P.', size=9, italic=True)

add_heading(doc, 'B\u1ea3ng 2. Hi\u1ec7u su\u1ea5t m\u00f4 h\u00ecnh d\u1ef1 \u0111o\u00e1n', 3)
add_table(doc,
    ['M\u00f4 h\u00ecnh', 'Feature Selection', 'Accuracy', 'AUC', 'Sensitivity', 'Precision', 'Specificity'],
    [['Random Forest', 'LASSO', '79%', '35%', '72%', '40%', '88%'],
     ['Random Forest', 'SelectKBest', '81%', '78%', '39%', '46%', '90%'],
     ['Random Forest', 'XGBoost', '78%', '82%', '64%', '43%', '81%'],
     ['ANN', 'LASSO', '80%', '74%', '31%', '41%', '90%'],
     ['ANN', 'SelectKBest', '82%', '81%', '44%', '51%', '90%'],
     ['ANN', 'XGBoost', '80%', '81%', '56%', '45%', '\u2014']],
    widths=[2.5, 2.5, 1.5, 1.5, 1.5, 1.5, 1.5])
add_p(doc, 'XGBoost + RF: AUC cao nh\u1ea5t (82%) + sensitivity t\u1ed1t nh\u1ea5t (64%). SE th\u1ea5p = hi\u1ec7u su\u1ea5t \u1ed5n \u0111\u1ecbnh.', size=9, italic=True)

# 4. THAO LUAN
add_page_ref(doc, '6\u20138', 'Frontiers in Public Health', '2025')
add_heading(doc, '4. TH\u1ea2O LU\u1eacN', 2)
add_p(doc, 'C\u00f4 \u0111\u01a1n l\u00e0 y\u1ebfu t\u1ed1 m\u1ea1nh nh\u1ea5t:', bold=True)
add_p(doc, 'C\u00f4 \u0111\u01a1n n\u1ed5i l\u00ean l\u00e0 y\u1ebfu t\u1ed1 \u1ea3nh h\u01b0\u1edfng nh\u1ea5t \u0111\u1ebfn GAD \u1edf VTN. VTN tr\u1ea3i qua c\u00f4 \u0111\u01a1n th\u01b0\u1eddng c\u00f3 \u00edt m\u1ea1ng l\u01b0\u1edbi x\u00e3 h\u1ed9i v\u00e0 h\u1ed7 tr\u1ee3 t\u00e2m l\u00fd, d\u1eabn \u0111\u1ebfn gia t\u0103ng b\u1ea5t \u1ed5n c\u1ea3m x\u00fac v\u00e0 nguy c\u01a1 GAD d\u00e0i h\u1ea1n (Ebesutani et al., 2015; Kim & Kim, 2009). ML x\u00e1c nh\u1eadn m\u1ea1nh m\u1ebd qua c\u1ea3 m\u00f4 h\u00ecnh tuy\u1ebfn t\u00ednh (LASSO) v\u00e0 phi tuy\u1ebfn (XGBoost) \u2014 m\u1ed1i quan h\u1ec7 c\u00f4 \u0111\u01a1n \u2192 GAD b\u1ec1n v\u1eefng b\u1ea5t k\u1ec3 ph\u01b0\u01a1ng ph\u00e1p. Ph\u00f9 h\u1ee3p Ireland 2024 (QT32, OGA b\u1ea3o v\u1ec7) v\u00e0 Islam 2025 (QT31, kh\u00f4ng b\u1ea1n th\u00e2n AOR=1,28).')
add_p(doc, 'Stress:', bold=True)
add_p(doc, 'Stress n\u1ed5i l\u00ean l\u00e0 y\u1ebfu t\u1ed1 th\u1ee9 2. Stress t\u1eeb quan h\u1ec7 b\u1ea1n b\u00e8, \u00e1p l\u1ef1c h\u1ecdc t\u1eadp, v\u00e0 c\u1ea3m gi\u00e1c b\u1ecb ph\u00e2n bi\u1ec7t \u0111\u1ed1i x\u1eed t\u1eeb gi\u00e1o vi\u00ean t\u0103ng \u0111\u00e1ng k\u1ec3 nguy c\u01a1 GAD (Lee & Cho, 2017). Ph\u00f9 h\u1ee3p Norway 2025 (QT21, b\u1ea5t m\u00e3n tr\u01b0\u1eddng gi\u1ea3i th\u00edch ch\u00ednh) v\u00e0 Wen 2020 (QT08, \u00e1p l\u1ef1c OR = 11,58).')
add_p(doc, 'B\u1ea1o l\u1ef1c:', bold=True)
add_p(doc, 'VTN t\u1eebng tr\u1ea3i qua b\u1ea1o l\u1ef1c th\u01b0\u1eddng ch\u1ecbu c\u1ea3 sang ch\u1ea5n th\u1ec3 ch\u1ea5t v\u00e0 th\u00e1ch th\u1ee9c SKTT. C\u00e1c NC tr\u01b0\u1edbc x\u00e1c nh\u1eadn m\u1ed1i li\u00ean quan m\u1ea1nh gi\u1eefa b\u1ea1o l\u1ef1c v\u00e0 t\u0103ng lo \u00e2u. C\u00e1c ch\u01b0\u01a1ng tr\u00ecnh h\u1ed7 tr\u1ee3 c\u1ea3m x\u00fac thi\u1ebft y\u1ebfu cho VTN n\u1ea1n nh\u00e2n b\u1ea1o l\u1ef1c. Ph\u00f9 h\u1ee3p Ng\u00f4 Anh Vinh 2024 (VN15, ACEs \u2192 lo \u00e2u Coef = 0,28).')
add_p(doc, 'H\u00e0nh vi s\u1ee9c kh\u1ecfe \u2014 C\u00d3 TH\u1ec2 CAN THI\u1ec6P:', bold=True)
add_p(doc, 'Ph\u1ee5c h\u1ed3i gi\u1ea5c ng\u1ee7, gi\u1ea5c ng\u1ee7 cu\u1ed1i tu\u1ea7n, h\u00fat thu\u1ed1c, v\u00e0 \u0103n th\u1ee9c \u0103n nhanh li\u00ean quan \u0111\u00e1ng k\u1ec3 v\u1edbi GAD. Ch\u1ea5t l\u01b0\u1ee3ng gi\u1ea5c ng\u1ee7 th\u1ea5p t\u01b0\u01a1ng quan v\u1edbi nguy c\u01a1 GAD cao h\u01a1n. Gi\u1ea5c ng\u1ee7 cu\u1ed1i tu\u1ea7n kh\u00f4ng \u0111\u1ee7 c\u00e0ng l\u00e0m tr\u1ea7m tr\u1ecdng v\u1ea5n \u0111\u1ec1 b\u1eb1ng c\u00e1ch ng\u0103n ph\u1ee5c h\u1ed3i t\u1eeb thi\u1ebfu ng\u1ee7 ng\u00e0y th\u01b0\u1eddng, d\u1eabn \u0111\u1ebfn thi\u1ebfu ng\u1ee7 m\u00e3n t\u00ednh \u2192 GAD. Ph\u00f9 h\u1ee3p Zhu 2025 (QT05, <5h AOR = 13,71).')
add_p(doc, 'Nicotine k\u00edch ho\u1ea1t h\u1ec7 th\u1ea7n kinh giao c\u1ea3m, t\u0103ng cortisol v\u00e0 norepinephrine, t\u1ea1o v\u00f2ng xo\u00e1y lo \u00e2u \u2192 h\u00fat thu\u1ed1c \u2192 lo \u00e2u.')
add_p(doc, 'N\u1eef gi\u1edbi v\u00e0 th\u1ee9c \u0103n nhanh:', bold=True)
add_p(doc, 'XGBoost ch\u1ecdn n\u1eef gi\u1edbi v\u00e0 \u0103n th\u1ee9c \u0103n nhanh h\u00e0ng ng\u00e0y l\u00e0 y\u1ebfu t\u1ed1 d\u1ef1 b\u00e1o. Nguy c\u01a1 lo \u00e2u cao \u1edf n\u1eef VTN \u0111\u01b0\u1ee3c ghi nh\u1eadn r\u1ed9ng r\u00e3i, b\u1ea5t k\u1ec3 \u0111i\u1ec1u ki\u1ec7n gia \u0111\u00ecnh, kinh t\u1ebf, hay h\u1ecdc t\u1eadp. \u00c1p l\u1ef1c x\u00e3 h\u1ed9i, lo ng\u1ea1i ngo\u1ea1i h\u00ecnh, k\u1ef3 v\u1ecdng h\u1ecdc t\u1eadp cao, v\u00e0 bi\u1ebfn \u0111\u1ed9ng hormone (estrogen) c\u00f3 th\u1ec3 t\u0103ng t\u00ednh d\u1ec5 b\u1ecb t\u1ed5n th\u01b0\u01a1ng. Ph\u00f9 h\u1ee3p Islam 2025 (QT31, n\u1eef AOR = 1,51 t\u1eeb 59 n\u01b0\u1edbc).')
add_p(doc, 'H\u1ea1n ch\u1ebf v\u00e0 h\u01b0\u1edbng t\u01b0\u01a1ng lai:', bold=True)
add_p(doc, 'NC c\u00f3 m\u1ed9t s\u1ed1 h\u1ea1n ch\u1ebf: (1) Thi\u1ebft k\u1ebf c\u1eaft ngang \u2014 kh\u00f4ng x\u00e1c l\u1eadp nh\u00e2n qu\u1ea3. (2) D\u1ef1a v\u00e0o d\u1eef li\u1ec7u t\u1ef1 b\u00e1o c\u00e1o KYRBS \u2014 c\u00f3 th\u1ec3 thi\u00ean l\u1ec7ch. (3) Kh\u00f4ng bao g\u1ed3m d\u1eef li\u1ec7u di truy\u1ec1n. NC t\u01b0\u01a1ng lai n\u00ean k\u1ebft h\u1ee3p d\u1eef li\u1ec7u di truy\u1ec1n \u0111\u1ec3 l\u00e0m r\u00f5 c\u01a1 ch\u1ebf GAD. Chi\u1ebfn l\u01b0\u1ee3c YTCC t\u1eadp trung c\u1ea3i thi\u1ec7n gi\u1ea5c ng\u1ee7, ph\u00f2ng h\u00fat thu\u1ed1c, v\u00e0 dinh d\u01b0\u1ee1ng c\u00e2n b\u1eb1ng c\u00f3 th\u1ec3 gi\u1ea3m GAD \u1edf VTN hi\u1ec7u qu\u1ea3.')

# 5. KET LUAN
add_heading(doc, '5. K\u1ebeT LU\u1eacN', 2)
add_p(doc, 'NC n\u00e0y x\u00e1c \u0111\u1ecbnh c\u00e1c y\u1ebfu t\u1ed1 ch\u00ednh li\u00ean quan GAD \u1edf VTN s\u1eed d\u1ee5ng ML tr\u00ean 213.820 VTN H\u00e0n Qu\u1ed1c. C\u00d4 \u0110\u01a0N l\u00e0 y\u1ebfu t\u1ed1 m\u1ea1nh nh\u1ea5t (LASSO 0,23; F=16.646; XGBoost 0,29), ti\u1ebfp theo stress, nh\u1eadn th\u1ee9c s\u1ee9c kh\u1ecfe k\u00e9m, ph\u1ee5c h\u1ed3i gi\u1ea5c ng\u1ee7 k\u00e9m, v\u00e0 b\u1ea1o l\u1ef1c. H\u00e0nh vi s\u1ee9c kh\u1ecfe (gi\u1ea5c ng\u1ee7, h\u00fat thu\u1ed1c, dinh d\u01b0\u1ee1ng) l\u00e0 y\u1ebfu t\u1ed1 C\u00d3 TH\u1ec2 CAN THI\u1ec6P. XGBoost + RF cho AUC 82%, sensitivity 64%.')

# TLTK
add_heading(doc, 'T\u00c0I LI\u1ec6U THAM KH\u1ea2O', 2)
for ref in [
    'Moon, Y. & Woo, H. (2025). Key risk factors of GAD in adolescents: ML study. Frontiers in Public Health, 13, 1504739.',
    'Spitzer, R.L. et al. (2006). A brief measure for assessing GAD: the GAD-7. Archives of Internal Medicine, 166(10), 1092\u20131097.',
    'Ebesutani, C. et al. (2015). The role of loneliness in the relationship between anxiety and depression. Psychology in the Schools, 52, 223\u2013234.',
    'Lee, J.H. & Cho, Y.T. (2017). Causal relationship of self-esteem, problem-solving, stress, and GAD in adolescents.',
    '(Xem \u0111\u1ea7y \u0111\u1ee7 41 TLTK trong b\u00e0i g\u1ed1c)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('GAD', 'Generalized Anxiety Disorder \u2014 R\u1ed1i lo\u1ea1n Lo \u00e2u T\u1ed5ng qu\u00e1t'),
    ('GAD-7', 'Generalized Anxiety Disorder 7-item Scale'),
    ('KYRBS', 'Korea Youth Risk Behavior Web-based Survey'),
    ('ML', 'Machine Learning \u2014 H\u1ecdc m\u00e1y'),
    ('LASSO', 'Least Absolute Shrinkage and Selection Operator'),
    ('XGBoost', 'eXtreme Gradient Boosting'),
    ('RF', 'Random Forest \u2014 R\u1eebng ng\u1eabu nhi\u00ean'),
    ('ANN', 'Artificial Neural Network \u2014 M\u1ea1ng N\u01a1-ron Nh\u00e2n t\u1ea1o'),
    ('SMOTE', 'Synthetic Minority Over-sampling Technique'),
    ('AUC', 'Area Under the Curve'),
    ('KDCA', 'Korea Disease Control and Prevention Agency'),
    ('DSM-5', 'Diagnostic and Statistical Manual of Mental Disorders, 5th Edition'),
    ('SKTT', 'S\u1ee9c kh\u1ecfe t\u00e2m th\u1ea7n'),
    ('VTN', 'V\u1ecb th\u00e0nh ni\u00ean'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N / CRITICAL REVIEW')
add_red(doc, '\u0110i\u1ec3m m\u1ea1nh:', bold=True)
for s in [
    'Frontiers Q1 IF\u22485,2. ML \u0111a ph\u01b0\u01a1ng ph\u00e1p \u2014 ph\u01b0\u01a1ng ph\u00e1p M\u1edaI trong NC GAD VTN.',
    'M\u1eabu C\u1ef0C L\u1edaN: 213.820 VTN \u2014 l\u1edbn nh\u1ea5t d\u00f9ng ML cho GAD. So: Hoa 2024 VN01 n=3.910.',
    'GAD-7 \u0111o lo \u00e2u T\u1ed4NG QU\u00c1T ri\u00eang \u2014 ch\u00ednh x\u00e1c h\u01a1n DASS-21.',
    '3 thu\u1eadt to\u00e1n ch\u1ecdn \u0111\u1eb7c tr\u01b0ng \u00d7 2 m\u00f4 h\u00ecnh = 6 t\u1ed5 h\u1ee3p \u2014 so s\u00e1nh \u0111a ph\u01b0\u01a1ng ph\u00e1p, gi\u1ea3m thi\u00ean l\u1ec7ch.',
    'KYRBS 2020\u20132023 \u2014 d\u1eef li\u1ec7u h\u1eadu COVID. THCS 53,8% + THPT 46,2% \u2014 bao ph\u1ee7 c\u1ea3 c\u1ea5p 2 v\u00e0 3.',
    'C\u00d4 \u0110\u01a0N h\u1ea1ng 1 \u2014 ph\u00e1t hi\u1ec7n quan tr\u1ecdng. Ph\u00f9 h\u1ee3p Ireland 2024 (QT32, OGA b\u1ea3o v\u1ec7), Islam 2025 (QT31, kh\u00f4ng b\u1ea1n th\u00e2n AOR=1,28).',
]:
    add_red(doc, '\u2022 ' + s)
add_red(doc, 'H\u1ea1n ch\u1ebf chi ti\u1ebft:', bold=True)
for s in [
    'Ch\u1ec9 H\u00e0n Qu\u1ed1c \u2014 kh\u00e1c VN. So: Tr\u1ea7n Th\u1ea3o Vi 2024 (VN21) c\u00f3 b\u1ed1i c\u1ea3nh VN nh\u01b0ng kh\u00f4ng ML.',
    'C\u1eaft ngang \u2014 ML x\u00e1c \u0111\u1ecbnh LI\u00caN QUAN, KH\u00d4NG nh\u00e2n qu\u1ea3. So: VN21 NC d\u1ecdc 3 n\u0103m.',
    'SMOTE c\u00f3 th\u1ec3 overfitting. Sensitivity 64% RF, 56% ANN = 36\u201344% VTN GAD b\u1ecb b\u1ecf s\u00f3t.',
    'GAD-7 \u2265 10 lo\u1ea1i b\u1ecf mild GAD (5\u20139). Hoa 2024 (VN01) d\u00f9ng \u2265 5: 40,6%.',
    'Kh\u00f4ng \u0111o: \u00e1p l\u1ef1c h\u1ecdc t\u1eadp (ESSA), screen time/MXH, ACEs \u2014 y\u1ebfu t\u1ed1 quan tr\u1ecdng theo Wen QT08, Norway QT21, Nature QT27, VN15.',
]:
    add_red(doc, '\u2022 ' + s)
add_red(doc, 'Kho\u1ea3ng tr\u1ed1ng NC / Gap:', bold=True)
for s in [
    'VN CH\u01afA C\u00d3 NC d\u00f9ng ML cho GAD \u1edf VTN. C\u00f3 th\u1ec3 \u00e1p d\u1ee5ng ML tr\u00ean V-NAMHS (VN02, n=5.996) ho\u1eb7c GSHS VN.',
    'So s\u00e1nh KYRBS H\u00e0n Qu\u1ed1c vs GSHS VN \u2014 c\u00f9ng ki\u1ec3u kh\u1ea3o s\u00e1t tr\u01b0\u1eddng.',
    'NC c\u00f4 \u0111\u01a1n \u1edf VTN VN \u2014 ch\u01b0a c\u00f3 NC ri\u00eang. Can thi\u1ec7p gi\u1ea3m c\u00f4 \u0111\u01a1n (mentor/OGA \u2014 Ireland QT32) c\u00f3 th\u1ec3 hi\u1ec7u qu\u1ea3 h\u01a1n can thi\u1ec7p stress.',
    'K\u1ebft h\u1ee3p ML + NC d\u1ecdc: ch\u1ea1y ML tr\u00ean Hue Cohort (VN21) \u2192 x\u00e1c \u0111\u1ecbnh y\u1ebfu t\u1ed1 D\u1ef0 B\u00c1O GAD.',
]:
    add_red(doc, '\u2022 ' + s)

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '37_Korea_GAD_ML_2025.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
checks = ['213.820', '17,9', '16.646', '0,29', '82%', '64%', '0,23', 'LASSO', 'XGBoost', 'KYRBS', 'Ebesutani']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
