# -*- coding: utf-8 -*-
"""Fix LA v3_4 -> v3_5: Chen 2023 errors.
- Para 335: 63.487 -> 63.205 (N analysis sample)
- Para 383: 'rối loạn stress' -> 'rối loạn lo âu hoặc trầm cảm' (Chen khong do stress)
27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_4_FixAbbrev_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_5_FixChen_27052026.docx')
RED = RGBColor(192, 0, 0)

REPLACEMENTS = [
    # Para 335: Chen N
    ('63.487 học sinh trung học cơ sở và trung học phổ thông',
     '63.205 học sinh trung học cơ sở và trung học phổ thông'),
    # Para 383: Chen IGD - 73.9% is distress (anxiety/depression), NOT stress
    ('trong số những học sinh nghiện game (internet gaming disorder – IGD) có 73,9% bị rối loạn stress',
     'trong số những học sinh nghiện game (internet gaming disorder – IGD) có 73,9% bị triệu chứng lo âu hoặc trầm cảm (mental distress)'),
]


def apply_repl(p, replacements, applied, idx):
    full = p.text
    if not full: return False
    matches = []
    for old, new in replacements:
        cursor = 0
        while True:
            i = full.find(old, cursor)
            if i == -1: break
            matches.append((i, i + len(old), old, new))
            cursor = i + len(old)
    if not matches: return False
    matches.sort(key=lambda x: (x[0], -(x[1]-x[0])))
    safe = []
    last = -1
    for m in matches:
        if m[0] >= last:
            safe.append(m); last = m[1]
    segments = []
    cur = 0
    for s, e, o, n in safe:
        if cur < s:
            segments.append((full[cur:s], False))
        segments.append((n, True))
        applied.append(f"para {idx}: '{o[:60]}' -> '{n[:60]}'")
        cur = e
    if cur < len(full):
        segments.append((full[cur:], False))
    for r in p.runs:
        r.text = ''
    for text, is_red in segments:
        if not text: continue
        r = p.add_run(text)
        r.font.name = 'Times New Roman'; r.font.size = Pt(13)
        if is_red:
            r.font.color.rgb = RED
            r.bold = True
    return True


print(f"=== Fix LA v3_4 -> v3_5 ===")
d = Document(SRC)
applied = []
for i, p in enumerate(d.paragraphs):
    apply_repl(p, REPLACEMENTS, applied, i)
d.save(OUT)
print(f"Saved: {OUT}")
print(f"Total: {len(applied)} fixes")
for a in applied:
    print(f"  ✓ {a}")
