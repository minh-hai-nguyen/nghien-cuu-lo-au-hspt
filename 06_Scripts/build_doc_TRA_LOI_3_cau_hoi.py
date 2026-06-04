"""DOC: TRA LOI 3 cau hoi cua thay (ngan gon, de doc):
1. Binh luan cua cac ban tro ly co on khong?
2. Tai sao khong quy ra %?
3. Cach quy ve %?

Format: CAU TRA LOI to xanh, ngan gon, de hieu cho thay.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/TRA_LOI_3_cau_hoi_banluan_va_phan_tram.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CÂU TRẢ LỜI 3 CÂU HỎI CỦA THẦY\n— Bàn luận của các bạn trợ lý + cách quy ra % —\n09/05/2026')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# CÂU HỎI 1
H('CÂU HỎI 1', level=2, color=NAVY)
blue(
    'Em xem bình luận của Bạn trợ lý của C.H có ổn không nhé? '
    'Thầy không hiểu biết gì nhiều về lĩnh vực này, ngoài những '
    'kiến thức mới học được từ Em trong thời gian qua.',
    italic=True
)

H('TRẢ LỜI 1 — Bình luận có cấu trúc ổn nhưng còn 6 thiếu sót cần sửa', level=2, color=GREEN)

para('Phần Bàn luận (đoạn 180–189 trong file v2) có ƯU và NHƯỢC:', bold=True)
para('')
para('ƯU điểm:', bold=True, color=GREEN)
para('• Cấu trúc theo từng yếu tố — dễ đọc.')
para('• Có dẫn chứng VN (Ngô Anh Vinh 2022, Lê T.P. Thanh 2024, Đinh Văn Tài 2021) và quốc tế (Carver 1997, Steare 2023, Li 2025, Fang 2022, Islam 2025, Feng 2024).')
para('• Phát hiện áp lực học tập là yếu tố mạnh nhất — đoạn 183.')
para('• Có giải thích β BPĐP DƯƠNG là maladaptive coping — đoạn 188.')
para('')
para('SÁU thiếu sót cần sửa:', bold=True, color=RED)
add_table(
    ['#', 'Thiếu sót', 'Đề xuất sửa'],
    [
        ['1', 'Sai năm "Nguyễn Thị Vân (2018)" — đoạn 183',
         'Sửa thành "Nguyễn Thị Vân (2020)" theo canonical_index'],
        ['2', 'BỎ QUA phát hiện CRITICAL về bắt nạt → lo âu CHIA LY',
         'Bổ sung: β BNHĐ → RLLACL = 0,376 MẠNH NHẤT trong 3 dạng RLLA — pattern KHÁC ALHT/TTr'],
        ['3', 'BỎ QUA cảnh báo fit indices KÉM của BPĐP',
         'Bổ sung: RMSEA 0,080–0,204 (vượt ngưỡng 0,06); CFI 0,865–0,911 (dưới 0,95) — KHÔNG đáng tin'],
        ['4', 'CHƯA nêu pattern ba tầng chênh lệch giới',
         'Bổ sung: RLLATQ (d=0,365), RLLAXH (d=0,370) — NGANG; RLLAC (d=0,03) KHÔNG'],
        ['5', 'Thiếu bảng so sánh cường độ chuẩn hóa các yếu tố',
         'Bổ sung bảng xếp hạng |β|: BPĐP 0,735 (fit kém) > ALHT 0,533 > TTr 0,457 > NĐT 0,400 > BNHĐ 0,276 > HTCM 0,234 > GBTH 0,155 > HTBB 0,044 NS'],
        ['6', 'KHÔNG quy ra % để dễ hình dung cỡ vấn đề',
         'Sẽ giải đáp ở Câu hỏi 2 và 3 dưới đây'],
    ]
)

# CÂU HỎI 2
H('CÂU HỎI 2', level=2, color=NAVY)
blue(
    'Tại sao số liệu thực trạng này lại không quy ra % như nhiều '
    'nghiên cứu khác?', italic=True
)

H('TRẢ LỜI 2 — Vì RCADS là thang đo MỨC ĐỘ, không phải chẩn đoán nhị phân', level=2, color=GREEN)

para(
    'Tác giả luận án sử dụng thang RCADS (Revised Children\'s '
    'Anxiety and Depression Scale; Chorpita và cộng sự, 2000) — '
    'là thang đo NỐI LIÊN TỤC (mức độ triệu chứng từ thấp đến cao), '
    'KHÁC với thang đo CHẨN ĐOÁN có/không như V-NAMHS hay Hoàng '
    'Trung Học sử dụng.'
)
para('')
add_table(
    ['Loại thang đo', 'Đặc trưng', 'Cho biết', 'Ví dụ'],
    [
        ['LIÊN TỤC (RCADS)',
         'Điểm số 0–100 hoặc 0–3 mỗi mục',
         'MỨC ĐỘ triệu chứng (cao/thấp)',
         'ĐTB = 51,43 → trung bình'],
        ['CHẨN ĐOÁN (DISC, CIDI)',
         'Có/không theo tiêu chí DSM-5',
         'TỶ LỆ % HS có rối loạn',
         'V-NAMHS: 18,45% có RLLA'],
    ]
)
para('')
para(
    'Để chuyển từ RCADS sang TỶ LỆ % cần một NGƯỠNG CẮT cụ '
    'thể — chương 3 luận án CHƯA xác định ngưỡng nào, do đó '
    'không thể quy ra % trực tiếp.', bold=True
)

# CÂU HỎI 3
H('CÂU HỎI 3', level=2, color=NAVY)
blue(
    'Có cách nào quy về % để có bức tranh thô về RLLA của HS '
    'không?', italic=True
)

H('TRẢ LỜI 3 — CÓ, ba cách. Em đã tính sẵn theo Cách 2 (z-score)', level=2, color=GREEN)

para('Ba cách quy về %:', bold=True)
add_table(
    ['Cách', 'Mô tả', 'Ưu', 'Nhược'],
    [
        ['1. Ngưỡng RCADS chuẩn',
         'T-score ≥ 65 (borderline) hoặc ≥ 70 (clinical)',
         'Theo chuẩn quốc tế Chorpita 2000',
         'Cần xác định norm — bài chưa nêu'],
        ['2. Z-score giả định normal',
         'z = (Ngưỡng − ĐTB) / SD; tra phân phối chuẩn',
         'Tính được ngay với data có sẵn',
         'GIẢ ĐỊNH normal — có thể sai'],
        ['3. Liên hệ tác giả xin data raw',
         'Đếm trực tiếp HS vượt ngưỡng',
         'CHÍNH XÁC NHẤT',
         'Cần liên hệ + chờ phản hồi'],
    ]
)

para('')
para('Em đã tính sẵn theo Cách 2 (ngưỡng 65/100, giả định normal):', bold=True)

H('Bức tranh thô — Tỷ lệ % HS có RLLA mức borderline (giả định normal)', level=3, color=NAVY)
add_table(
    ['Loại RLLA', 'ĐTB', 'SD', 'z', '% TỔNG', '% Nam', '% Nữ'],
    [
        ['RLLATQ (lan tỏa)', '55,82', '22,04', '0,42', '≈ 33,85%', '≈ 26,88%', '≈ 40,11%'],
        ['RLLAXH (xã hội)', '48,41', '25,77', '0,64', '≈ 25,98%', '≈ 19,25%', '≈ 32,06%'],
        ['RLLAC (chia ly)', '25,06', '24,30', '1,64', '≈ 5,01%', '−', '−'],
    ]
)
para('')

para('So sánh với các nghiên cứu Việt Nam khác:', bold=True)
add_table(
    ['Nghiên cứu', 'Mẫu', 'Tỷ lệ %', 'Ghi chú'],
    [
        ['V-NAMHS (VN002, 2022)', '5.996 HS 10–17', '18,45%', 'Chẩn đoán DSM-5 (chính thức)'],
        ['Hoàng Trung Học (VN014, 2025)', '8.473 HS', '41,5% / 25,4%', 'Sàng lọc DASS-21 trong/sau COVID'],
        ['Chương 3 luận án — RLLATQ', '1.352 HS THCS', '≈ 33,85%', 'Ước lượng z-score'],
        ['Chương 3 luận án — RLLAXH', '1.352 HS THCS', '≈ 25,98%', 'Ước lượng z-score'],
        ['Chương 3 luận án — RLLAC', '1.352 HS THCS', '≈ 5,01%', 'Ước lượng z-score'],
    ]
)
para('')
para(
    'Tỷ lệ ước lượng RLLAXH (25,98%) GẦN với V-NAMHS chẩn '
    'đoán (18,45%) và Hoàng Trung Học sàng lọc sau COVID '
    '(25,4%). RLLATQ (33,85%) cao hơn — phù hợp với đặc thù '
    'lo âu lan tỏa phổ biến hơn ở thanh thiếu niên.', bold=True
)

# Cảnh báo
H('CẢNH BÁO QUAN TRỌNG', level=2, color=RED)
para(
    'Tỷ lệ % ở Cách 2 là ƯỚC LƯỢNG GẦN ĐÚNG dựa trên giả định '
    'phân phối CHUẨN. Trong y văn, dữ liệu RLLA thường LỆCH '
    'PHẢI (đa số học sinh không có triệu chứng) — tỷ lệ thực '
    'tế có thể THẤP HƠN ước lượng. Trước khi đưa số % vào báo '
    'cáo CTH, NÊN: (a) liên hệ tác giả luận án xin DATA RAW; '
    '(b) hoặc ghi chú rõ "ước lượng giả định normal" trong báo '
    'cáo.', bold=True, color=RED
)

# Khuyến nghị
H('KHUYẾN NGHỊ TỔNG cho thầy', level=2, color=NAVY)
blue('Ba việc cần làm với phần Bàn luận 3.6:', bold=True)
blue(
    '(1) Sửa 6 thiếu sót đã liệt kê — đặc biệt sửa năm Nguyễn '
    'Thị Vân 2020 và bổ sung 2 phát hiện đặc biệt (BNHĐ → '
    'RLLACL; cảnh báo BPĐP fit kém).'
)
blue(
    '(2) Bổ sung BẢNG % vào báo cáo (theo Cách 2 hoặc Cách 3) '
    'để người đọc có "bức tranh thô" về quy mô vấn đề. Ưu '
    'tiên Cách 3 (xin data raw từ tác giả) cho độ chính xác.'
)
blue(
    '(3) Trong khi chờ data raw, dùng số % ước lượng từ Cách 2: '
    '~34% HS có lo âu lan tỏa borderline; ~26% có lo âu xã '
    'hội; ~5% có lo âu chia ly. Ghi chú rõ "ước lượng giả '
    'định normal".'
)

# File chi tiết
H('Tham chiếu', level=2, color=NAVY)
para(
    'Để xem PHÉP TOÁN chi tiết và lập luận đầy đủ, xem file:'
)
para(
    '01_Bao-cao/Danh_gia_banluan_3_6_va_cach_quy_ve_phan_tram.docx '
    '(41 KB) — bao gồm 5 phần: Đánh giá Bàn luận; Lý do không '
    'có %; 3 cách quy về %; So sánh với prevalence quốc tế; '
    'Đề xuất tổng.', italic=True
)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
