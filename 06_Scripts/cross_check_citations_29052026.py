# -*- coding: utf-8 -*-
"""Cross-check citations trong cac doc viet vs PDF catalog.
Trich tat ca Author (Year) trong each doc, match voi catalog,
flag mismatches va missing PDFs.
29/05/2026."""
import os, sys, io, json, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Load PDF catalog
with open(os.path.join(ROOT, '06_Scripts', 'pdf_catalog_29052026.json'),
          'r', encoding='utf-8') as f:
    catalog = json.load(f)

print(f'Loaded {len(catalog["pdfs"])} PDFs')

# Build lookup: filename surface forms → metadata
pdf_index = {}
for p in catalog['pdfs']:
    fn = p['filename'].lower()
    pdf_index[fn] = p
    # Also index by filename without prefix
    base = fn.replace('.pdf', '')
    pdf_index[base] = p

# Build year→PDFs index
year_to_pdfs = {}
for p in catalog['pdfs']:
    y = p.get('extracted_year')
    if y:
        year_to_pdfs.setdefault(y, []).append(p)


def find_pdfs_by_author_year(author, year):
    """Tim PDF khop author + year tu catalog."""
    if year not in year_to_pdfs:
        return []
    candidates = year_to_pdfs[year]
    author_l = author.lower()
    matches = []
    for p in candidates:
        fn = p['filename'].lower()
        title = (p['extracted_title'] or '').lower()
        meta = (p['meta_author'] or '').lower()
        first = p['first_500'].lower()
        if author_l in fn or author_l in title or author_l in meta or author_l in first:
            matches.append(p)
    return matches


# Target docs to extract citations from
DOCS = {
    'TT_VN': 'Luận án TS/TomTatLA_v2_VERIFIED_29052026.docx',
    'TT_EN': 'Luận án TS/TomTatLA_EN_v1_29052026.docx',
    'TY': 'Luận án TS/TrichYeuLA_CongThiHang_v2_29052026.docx',
    'BaiA': 'bai-bao-Q1/BaiA_JAD_OUTLINE_v1_30052026.docx',
    'BaiD': 'bai-bao-Q1/BaiD_StressHealth_OUTLINE_v1_30052026.docx',
    'LA': 'Luận án TS/01_LuanAn_v4_ChuanTrinhBay_28052026.docx',
}


def extract_citations(text):
    """Extract Author + Year patterns. Returns list of (author, year, full_match)."""
    cites = []
    # Pattern: "Last Name (Year)" or "Last Name et al. (Year)" or "Last Name & X (Year)"
    # Also: "Author, Year" in parens
    patterns = [
        r'([A-Z][a-zA-Zà-ỹ]+(?:\s+(?:&|and|và)\s+[A-Z][a-zA-Zà-ỹ]+)?(?:\s+et\s+al\.?)?)\s*\(([12]\d{3}[a-z]?)\)',
        r'\(([A-Z][a-zA-Zà-ỹ]+(?:\s+(?:&|and|và)\s+[A-Z][a-zA-Zà-ỹ]+)?(?:\s+et\s+al\.?)?)[,;]?\s*([12]\d{3}[a-z]?)\)',
    ]
    for pat in patterns:
        for m in re.finditer(pat, text):
            author = m.group(1).strip()
            year = m.group(2)
            cites.append((author, year, m.group(0)))
    # dedupe
    return list(set(cites))


# Extract citations from each doc
all_cites = {}
for key, path in DOCS.items():
    full = os.path.join(ROOT, path)
    if not os.path.exists(full):
        continue
    d = Document(full)
    text = '\n'.join(p.text for p in d.paragraphs)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                text += '\n' + c.text
    cites = extract_citations(text)
    all_cites[key] = cites
    print(f'{key}: {len(cites)} unique citations')


# Cross-check each citation against PDF catalog
print('\n=== CROSS-CHECK ===')
total_cites = 0
matched = 0
unmatched = []
year_mismatches = []

for fkey, cites in all_cites.items():
    for author, year, full_match in cites:
        total_cites += 1
        year_int = int(re.match(r'(\d{4})', year).group(1))

        # Skip common non-author words
        skip_words = ['Bảng', 'Hình', 'Chương', 'Mục', 'NCS', 'TS', 'PGS', 'GS',
                      'Table', 'Figure', 'Chapter', 'Section',
                      'DSM', 'WHO', 'GBD', 'COVID', 'COVID-19']
        if any(s in author for s in skip_words):
            continue

        matches = find_pdfs_by_author_year(author.lower().replace(' et al.', '').replace(' và cs', '').split()[0], year_int)
        if not matches:
            # Try year ± 1 (publication year shift)
            for y in [year_int - 1, year_int + 1]:
                if y in year_to_pdfs:
                    matches = find_pdfs_by_author_year(
                        author.lower().replace(' et al.', '').replace(' và cs', '').split()[0], y)
                    if matches:
                        year_mismatches.append(
                            (fkey, author, year, [p['filename'] for p in matches][:2]))
                        break

        if matches:
            matched += 1
        else:
            unmatched.append((fkey, author, year, full_match))


print(f'\nTotal citations: {total_cites}')
print(f'Matched: {matched}')
print(f'Unmatched: {len(unmatched)}')
print(f'Year mismatches: {len(year_mismatches)}')

# Save report
report_path = os.path.join(ROOT, '06_Scripts', 'citation_crosscheck_29052026.json')
with open(report_path, 'w', encoding='utf-8') as f:
    json.dump({
        'total_citations': total_cites,
        'matched': matched,
        'unmatched': [{'file': u[0], 'author': u[1], 'year': u[2], 'context': u[3]}
                      for u in unmatched],
        'year_mismatches': [{'file': m[0], 'author': m[1], 'cited_year': m[2],
                             'pdf_matches': m[3]} for m in year_mismatches],
    }, f, ensure_ascii=False, indent=2)
print(f'\nSaved: {report_path}')

# Top 30 unmatched citations (potential errors)
print('\n=== TOP UNMATCHED CITATIONS ===')
unmatched_count = {}
for f, a, y, _ in unmatched:
    key = f'{a} ({y})'
    unmatched_count[key] = unmatched_count.get(key, 0) + 1
for k, c in sorted(unmatched_count.items(), key=lambda x: -x[1])[:30]:
    print(f'  {c}× {k}')
