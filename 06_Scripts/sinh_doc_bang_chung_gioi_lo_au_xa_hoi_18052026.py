# -*- coding: utf-8 -*-
"""Sinh doc tra loi thay: bang chung quoc te ve khac biet gioi trong lo au xa hoi.
18/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Tra_loi_bang_chung_gioi_lo_au_xa_hoi_18052026.docx')
PAGE_W = 15.5
RED = RGBColor(0xCC, 0, 0)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)

def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(doc, text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def bullet(doc, text, bold_prefix=None, size=12, color=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = True
        if color is not None: r.font.color.rgb = color
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    if color is not None: r.font.color.rgb = color
    return p

def add_hyperlink(paragraph, url, text, size=11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hyperlink = OxmlElement('w:hyperlink'); hyperlink.set(qn('r:id'), r_id)
    new_run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')
    rFonts = OxmlElement('w:rFonts')
    rFonts.set(qn('w:ascii'), 'Times New Roman'); rFonts.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rFonts)
    color = OxmlElement('w:color'); color.set(qn('w:val'), '0563C1'); rPr.append(color)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
    new_run.append(rPr)
    t = OxmlElement('w:t'); t.text = text; new_run.append(t)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)

def ref_item(doc, text, doi_label=None, doi_url=None):
    p = doc.add_paragraph(style='List Bullet')
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    if doi_label and doi_url:
        sep = p.add_run('  '); sep.font.name = 'Times New Roman'; sep.font.size = Pt(11)
        add_hyperlink(p, doi_url, doi_label, size=11)

def table(doc, headers, rows, widths, hdr_size=10, body_size=9.5):
    assert abs(sum(widths) - PAGE_W) < 0.1, sum(widths)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(hdr_size)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(body_size)
    return t

# ============================================================
doc = make_doc()

ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('TRẢ LỜI: BẰNG CHỨNG QUỐC TẾ VỀ KHÁC BIỆT GIỚI\n'
                'TRONG LO ÂU XÃ HỘI — NỮ CÓ CAO HƠN NAM KHÔNG?')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Tài liệu trả lời câu hỏi — Ngày 18/05/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

P(doc, 'Câu hỏi: Với lo âu lan tỏa (Generalized Anxiety Disorder — GAD) thì đã rõ nữ có mức '
       'độ cao hơn nam. Còn với lo âu xã hội (Social Anxiety Disorder — SAD), có nhiều bằng '
       'chứng quốc tế cho thấy nữ cao hơn nam không?', italic=True, size=11)

# ---- 1. Tra loi ngan gon ----
H(doc, '1. Trả lời ngắn gọn', 1)
P(doc, 'Thưa thầy, CÓ bằng chứng quốc tế cho thấy nữ có lo âu xã hội cao hơn nam — nhưng để '
       'chính xác cần tách thành hai khía cạnh, vì độ chắc chắn của hai khía cạnh này khác '
       'nhau:')
bullet(doc, 'Bằng chứng MẠNH và rất nhất quán — nữ cao hơn nam, lặp lại ở nhiều quốc gia.',
       bold_prefix='Về mức độ triệu chứng (đo bằng thang tự báo cáo): ')
bullet(doc, 'Bằng chứng nghiêng về phía nữ cao hơn, nhưng KHÔNG tuyệt đối. Lo âu xã hội là '
            'rối loạn lo âu có khác biệt giới HẸP NHẤT và gây tranh luận nhiều nhất; có '
            'nghiên cứu lớn không tìm thấy khác biệt nào.',
       bold_prefix='Về tỷ lệ được chẩn đoán trong cộng đồng: ')
P(doc, 'Vì vậy, sự thận trọng của thầy là có cơ sở: lo âu xã hội đúng là không có khác biệt '
       'giới "rõ như ban ngày" giống lo âu lan tỏa. Tuy nhiên, riêng ở lứa VỊ THÀNH NIÊN — '
       'đối tượng của bản thảo bai2 — bằng chứng nữ cao hơn nam lại nhất quán hơn so với ở '
       'người trưởng thành.')

# ---- 2. Bang chung ----
H(doc, '2. Bằng chứng theo từng loại dữ liệu', 1)

H(doc, '2.1. Mức độ triệu chứng đo bằng thang tự báo cáo', 2)
P(doc, 'Đây là khía cạnh có bằng chứng mạnh nhất. Nghiên cứu của Caballo và cộng sự (2014) — '
       'một nghiên cứu xuyên văn hóa trên 18 quốc gia (16 nước Mỹ Latinh cùng Tây Ban Nha và '
       'Bồ Đào Nha) — đo lo âu xã hội trên mẫu cộng đồng gồm 17.672 nữ và 13.440 nam (cùng '
       '601 bệnh nhân lo âu xã hội), sử dụng Bảng hỏi Lo âu Xã hội cho Người trưởng thành '
       '(Social Anxiety Questionnaire for Adults — SAQ-A). Kết quả: nữ báo cáo lo âu xã hội '
       'cao hơn nam ở 88,67% các tình huống xã hội được khảo sát; 15 trong 18 quốc gia ghi '
       'nhận khác biệt giới có ý nghĩa thống kê theo hướng nữ cao hơn. Khác biệt này được mô '
       'tả là "nhỏ nhưng có ý nghĩa". Vì lặp lại nhất quán xuyên nhiều nền văn hóa, đây là '
       'bằng chứng đáng tin cậy nhất cho nhận định nữ lo âu xã hội nhiều hơn nam.')

H(doc, '2.2. Tỷ lệ được chẩn đoán trong cộng đồng', 2)
P(doc, 'Ở khía cạnh tỷ lệ mắc (được chẩn đoán), bức tranh kém nhất quán hơn:')
bullet(doc, 'Sổ tay Chẩn đoán và Thống kê các Rối loạn Tâm thần phiên bản 5 (DSM-5, Hiệp hội '
            'Tâm thần học Hoa Kỳ, 2013) nêu rõ: tỷ lệ lo âu xã hội cao hơn ở nữ, và khác biệt '
            'này rõ rệt hơn ở lứa vị thành niên.')
bullet(doc, 'Khảo sát NESARC của Hoa Kỳ (phân tích của Xu và cộng sự, 2012): tỷ lệ mắc lo âu '
            'xã hội trọn đời là 5,67% ở nữ so với 4,20% ở nam — nữ cao hơn, tỷ số khoảng '
            '1,35:1.')
bullet(doc, 'NGƯỢC LẠI: McLean và cộng sự (2011) phân tích dữ liệu CPES trên 20.013 người '
            'trưởng thành Hoa Kỳ, phát hiện nữ có tỷ lệ mắc cao hơn ở MỌI rối loạn lo âu được '
            'khảo sát — NGOẠI TRỪ lo âu xã hội: riêng lo âu xã hội KHÔNG có khác biệt giới về '
            'tỷ lệ mắc.')
P(doc, 'Như vậy, ở khía cạnh chẩn đoán, đa số nghiên cứu nghiêng về phía nữ cao hơn, nhưng '
       'không nhất quán tuyệt đối — lo âu xã hội nổi lên như một "ngoại lệ" trong nhóm rối '
       'loạn lo âu.', color=RED)

H(doc, '2.3. Riêng ở lứa vị thành niên', 2)
P(doc, 'Đối với vị thành niên, bằng chứng nữ cao hơn nam nhất quán hơn. DSM-5 ghi nhận khác '
       'biệt nữ – nam trong lo âu xã hội rõ rệt hơn ở lứa tuổi này. Khảo sát NCS-A (Khảo sát '
       'Bổ sung Vị thành niên thuộc chương trình National Comorbidity Survey, Merikangas và '
       'cộng sự, 2010) trên vị thành niên Hoa Kỳ cũng cho thấy nữ vị thành niên có tỷ lệ lo '
       'âu xã hội cao hơn nam (ước tính khoảng 11% so với 7%). Một số nghiên cứu theo dõi dọc '
       'còn ghi nhận triệu chứng lo âu xã hội ở nữ có xu hướng tăng dần qua tuổi vị thành '
       'niên trong khi ở nam lại giảm; ngoài ra, dậy thì sớm gắn với tăng triệu chứng lo âu '
       'xã hội nhưng chỉ ở nữ. Đây là điểm quan trọng cho bản thảo bai2 (vốn viết về học sinh '
       'trung học cơ sở): ở nhóm tuổi này, nhận định nữ cao hơn nam có cơ sở vững hơn so với '
       'khi nói chung về người trưởng thành.')

H(doc, '2.4. Khác biệt giữa mẫu cộng đồng và mẫu lâm sàng', 2)
P(doc, 'Bài tổng quan của Asher, Asnaani và Aderka (2017) rà soát 8 phương diện khác biệt '
       'giới trong lo âu xã hội (tỷ lệ mắc, biểu hiện lâm sàng, suy giảm chức năng, bệnh đồng '
       'diễn, diễn tiến, tìm kiếm điều trị, kích hoạt sinh lý, hệ oxytocin). Kết luận: trong '
       'cộng đồng, nữ mắc nhiều hơn và có biểu hiện nặng hơn (nhiều nỗi sợ xã hội hơn, mức '
       'khó chịu chủ quan cao hơn). Tuy nhiên, trong mẫu lâm sàng — tức nhóm bệnh nhân thực '
       'sự đến khám — tỷ lệ nam và nữ lại cân bằng hơn, vì nam mắc lo âu xã hội có xu hướng '
       'đi điều trị nhiều hơn. Một cách lý giải là áp lực vai trò giới: xã hội kỳ vọng nam '
       'giới phải quyết đoán, hoạt ngôn nơi công cộng, nên lo âu xã hội gây tổn hại rõ rệt '
       'tới đời sống của họ và thúc đẩy họ tìm trợ giúp.')

# ---- 3. Bang tom tat ----
H(doc, '3. Bảng tóm tắt bằng chứng', 1)
table(doc,
    ['Loại bằng chứng / nghiên cứu', 'Phát hiện chính về khác biệt giới', 'Nguồn'],
    [
        ['Thang đo tự báo cáo, mẫu cộng đồng xuyên 18 quốc gia (hơn 31.000 người)',
         'Nữ lo âu xã hội cao hơn nam ở 88,67% tình huống xã hội; 15/18 quốc gia có khác '
         'biệt có ý nghĩa thống kê (nữ cao hơn).',
         'Caballo và cs. (2014)'],
        ['Tỷ lệ mắc trọn đời — khảo sát NESARC (Hoa Kỳ)',
         'Nữ 5,67% — nam 4,20% (nữ cao hơn, tỷ số khoảng 1,35:1).',
         'Xu và cs. (2012)'],
        ['Tỷ lệ mắc — dữ liệu CPES, 20.013 người trưởng thành Hoa Kỳ',
         'KHÔNG có khác biệt giới ở lo âu xã hội — dù nữ cao hơn ở mọi rối loạn lo âu khác.',
         'McLean và cs. (2011)'],
        ['Tổng quan hệ thống 8 phương diện của lo âu xã hội',
         'Trong cộng đồng nữ mắc nhiều hơn và nặng hơn; trong mẫu lâm sàng nam đi điều trị '
         'nhiều hơn nên tỷ lệ cân bằng hơn.',
         'Asher, Asnaani & Aderka (2017)'],
        ['Vị thành niên — chuẩn phân loại & khảo sát NCS-A',
         'Lo âu xã hội phổ biến hơn ở nữ; khác biệt nữ – nam rõ rệt hơn ở lứa vị thành niên.',
         'DSM-5 (APA, 2013); Merikangas và cs. (2010)'],
    ],
    widths=[4.8, 7.0, 3.7])

# ---- 4. So sanh GAD ----
H(doc, '4. So sánh với lo âu lan tỏa', 1)
P(doc, 'Sự khác nhau về độ chắc chắn giữa hai rối loạn có thể tóm tắt như sau. Với lo âu lan '
       'tỏa, tỷ số nữ : nam khá ổn định ở mức khoảng 2 : 1 qua nhiều nghiên cứu — đây là lý '
       'do thầy thấy "đã rõ". Với lo âu xã hội, khác biệt giới hẹp hơn rõ rệt: dao động từ '
       '"không có khác biệt" (dữ liệu CPES của McLean và cộng sự) đến khoảng 1,35 : 1 (khảo '
       'sát NESARC). Nói cách khác, cảm nhận của thầy — "lo âu lan tỏa thì rõ, lo âu xã hội '
       'thì chưa chắc" — phản ánh đúng thực trạng tài liệu khoa học, chứ không phải do thiếu '
       'bằng chứng.')

# ---- 5. Luu y ----
H(doc, '5. Lưu ý khi viết và trình bày', 1)
bullet(doc, 'Khi khẳng định nữ cao hơn nam, nên ghi rõ đang nói về "mức độ triệu chứng" hay '
            '"tỷ lệ được chẩn đoán" — hai khía cạnh có độ chắc chắn khác nhau.',
       bold_prefix='Tách hai khía cạnh. ')
bullet(doc, 'Nên ghi rõ "trong cộng đồng" hay "trong mẫu lâm sàng", vì khác biệt giới đảo '
            'chiều giữa hai loại mẫu này.', bold_prefix='Ghi rõ loại mẫu. ')
bullet(doc, 'Không nên mượn tỷ số 2 : 1 của lo âu lan tỏa để áp cho lo âu xã hội — sẽ phóng '
            'đại khác biệt.', bold_prefix='Không dùng chung tỷ số. ')
bullet(doc, 'Với bài viết về học sinh trung học cơ sở, có thể tự tin dùng nhận định nữ cao '
            'hơn nam, vì ở lứa vị thành niên bằng chứng vững hơn — nhưng vẫn nên dẫn nguồn cụ '
            'thể.', bold_prefix='Riêng vị thành niên. ')

# ---- 6. Y nghia VN ----
H(doc, '6. Ý nghĩa cho bối cảnh Việt Nam', 1)
P(doc, 'Phần lớn bằng chứng trên đến từ Hoa Kỳ và các nước phương Tây / Mỹ Latinh. Để có dẫn '
       'liệu nội địa, nên đối chiếu số liệu theo giới của Khảo sát Sức khỏe Tâm thần Vị thành '
       'niên Việt Nam (V-NAMHS) — khảo sát này có đo ám ảnh sợ xã hội ở trẻ 10–17 tuổi, nếu '
       'có bảng phân tách theo giới thì đây sẽ là bằng chứng Việt Nam trực tiếp nhất.')
P(doc, 'Về thực hành can thiệp học đường: nữ vị thành niên có thể là nhóm cần ưu tiên sàng '
       'lọc lo âu xã hội. Tuy nhiên cũng cần lưu ý chiều ngược lại — nam vị thành niên có xu '
       'hướng ít bộc lộ và ít chủ động tìm trợ giúp, nên dễ bị bỏ sót khi sàng lọc; công cụ '
       'sàng lọc và cách tiếp cận nên tính đến đặc điểm này để không "lọt lưới" học sinh nam.')

# ---- 7. Tham khao ----
H(doc, '7. Tham khảo & truy vết', 1)
P(doc, 'Nguồn học thuật (đã đối chiếu để xác minh):', bold=True, size=11)
ref_item(doc,
    'Asher, M., Asnaani, A., & Aderka, I. M. (2017). "Gender differences in social anxiety '
    'disorder: A review." Clinical Psychology Review, 56, 1–12. PMID: 28578248. — Tổng quan '
    'cốt lõi về khác biệt giới trong lo âu xã hội.',
    'doi:10.1016/j.cpr.2017.05.004', 'https://doi.org/10.1016/j.cpr.2017.05.004')
ref_item(doc,
    'McLean, C. P., Asnaani, A., Litz, B. T., & Hofmann, S. G. (2011). "Gender differences '
    'in anxiety disorders: Prevalence, course of illness, comorbidity and burden of '
    'illness." Journal of Psychiatric Research, 45(8), 1027–1035. — Dữ liệu CPES (N = '
    '20.013); lo âu xã hội KHÔNG có khác biệt giới về tỷ lệ mắc.',
    'doi:10.1016/j.jpsychires.2011.03.006',
    'https://doi.org/10.1016/j.jpsychires.2011.03.006')
ref_item(doc,
    'Caballo, V. E., Salazar, I. C., Irurtia, M. J., Arias, B., Hofmann, S. G., & CISO-A '
    'Research Team (2014). "Differences in social anxiety between men and women across 18 '
    'countries." Personality and Individual Differences, 64, 35–40. PMID: 24976665. — Bằng '
    'chứng xuyên văn hóa: nữ lo âu xã hội cao hơn nam.')
ref_item(doc,
    'Xu, Y., Schneier, F., Heimberg, R. G., Princisvalle, K., Liebowitz, M. R., Wang, S., & '
    'Blanco, C. (2012). "Gender differences in social anxiety disorder: Results from the '
    'National Epidemiologic Sample on Alcohol and Related Conditions." Journal of Anxiety '
    'Disorders, 26(1), 12–19. PMID: 21903358. — Khảo sát NESARC.')
ref_item(doc,
    'Merikangas, K. R., He, J. P., Burstein, M., và cộng sự (2010). "Lifetime prevalence of '
    'mental disorders in U.S. adolescents: Results from the National Comorbidity Survey '
    'Replication–Adolescent Supplement (NCS-A)." Journal of the American Academy of Child & '
    'Adolescent Psychiatry, 49(10), 980–989. — Dữ liệu tỷ lệ rối loạn tâm thần ở vị thành '
    'niên Hoa Kỳ.')
ref_item(doc,
    'American Psychiatric Association (2013). Diagnostic and Statistical Manual of Mental '
    'Disorders (5th ed. — DSM-5). Arlington, VA: American Psychiatric Publishing. — Nêu rõ '
    'lo âu xã hội phổ biến hơn ở nữ, khác biệt rõ hơn ở vị thành niên.')

P(doc, 'Truy vết nội bộ:', bold=True, size=11)
bullet(doc, 'Câu hỏi phát sinh trong quá trình rà soát bản thảo bai2 '
            '(06_Scripts/build_bai2_KHGDVN.py) — bài viết về lo âu ở học sinh trung học cơ '
            'sở. Tài liệu này bổ sung cơ sở cho phần bàn về yếu tố giới.', size=11)
bullet(doc, 'Gợi ý đối chiếu nội địa: báo cáo Khảo sát Sức khỏe Tâm thần Vị thành niên Việt '
            'Nam (V-NAMHS) — nếu có bảng số liệu ám ảnh sợ xã hội phân tách theo giới.',
       size=11)

P(doc, 'Lưu ý xác minh: các con số và phát hiện trong tài liệu này được đối chiếu qua '
       'PubMed và trang tạp chí gốc. Riêng số liệu vị thành niên (khoảng 11% nữ so với 7% '
       'nam) là ước tính từ dữ liệu NCS-A; khi trích dẫn chính thức nên tra trực tiếp bảng '
       'số liệu gốc của NCS-A để lấy con số chính xác theo từng nhóm tuổi.', italic=True,
  size=11)

doc.save(OUT)
print('DA LUU:', OUT)
print('So bang:', len(doc.tables), '| So doan:', len(doc.paragraphs))
