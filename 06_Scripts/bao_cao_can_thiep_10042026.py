# -*- coding: utf-8 -*-
"""
Tạo BÁO CÁO BIỆN PHÁP HỖ TRỢ/CAN THIỆP TÂM LÝ CHO VTN RỐI LOẠN LO ÂU
Phân thành 3 mục: Việt Nam | Châu Á | Âu-Úc-Mỹ
Dữ liệu từ 14+ bài can thiệp đã có trong hệ thống (B1-B11 + QT28-29 + Trần Nguyễn Ngọc 2018)
"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Bao cao Can thiep tam ly RLLA VTN - 10042026.docx')

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    tcW.append(w)

def H(text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def P(text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color: r.font.color.rgb = color
    return p

def table(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
    return t

# ==========================================================================
# TIÊU ĐỀ
# ==========================================================================
H('BÁO CÁO TỔNG HỢP', level=1)
H('BIỆN PHÁP HỖ TRỢ TÂM LÝ / CAN THIỆP TÂM LÝ CHO TRẺ VỊ THÀNH NIÊN BỊ RỐI LOẠN LO ÂU', level=1)
P('Cập nhật: 10/04/2026 | Tổng hợp 15 công trình từ hệ thống 58 bài NC', italic=True)
P('Phân theo địa lý: (1) Việt Nam — 1 bài, (2) Châu Á (ngoài VN) — 7 bài, (3) Âu - Úc - Mỹ — 7 bài', italic=True)

# ==========================================================================
# TÓM LƯỢC ĐIỀU HÀNH
# ==========================================================================
H('TÓM LƯỢC ĐIỀU HÀNH', level=2)
P('Báo cáo này tổng hợp 15 công trình về can thiệp/hỗ trợ tâm lý cho RLLA ở trẻ VTN, được phân theo 3 vùng địa lý. Bằng chứng cho thấy:', bold=True)
P('• CBT (Cognitive Behavioral Therapy) là phương pháp được chứng minh hiệu quả nhất xuyên suốt 3 vùng, với SUCRA 0,66 (BMC NMA 2025) và 80,7% đáp ứng (CAMS Walkup 2008 NEJM).')
P('• iCBT (internet-based CBT) cho lo âu xã hội (SAD) đứng hạng 1 trong network meta-analysis 30 RCT (Xian 2024, SUCRA 71,2%) — phù hợp cả Châu Á và Phương Tây.')
P('• Mobile CBT mạnh cho TRẦM CẢM (87,5% NC dương tính) nhưng YẾU cho LO ÂU (33% NC dương tính) — Qiaochu 2025 SR 9 RCT.')
P('• Universal CBT (CBT phổ quát tại trường) cho hiệu quả NHỎ với chất lượng nền THẤP (Zhang 2026 Bayesian MA 31 RCT n=19.865) — TARGETED CBT (cho HS có triệu chứng) hứa hẹn hơn.')
P('• KHOẢNG TRỐNG LỚN ở Việt Nam: 0 RCT can thiệp tâm lý cho RLLA ở VTN VN trong hệ thống. Bài VN duy nhất là Trần Nguyễn Ngọc 2018 (Bệnh viện Bạch Mai) — nhưng đối tượng là NGƯỜI LỚN nội trú, không phải VTN.')
P('• Khu vực Đông Á + Thái Bình Dương LMIC còn thiếu can thiệp CỘNG ĐỒNG + GIA ĐÌNH + dài hạn (Menon 2025 Scoping 69 NC).')

# ==========================================================================
# I. VIỆT NAM
# ==========================================================================
doc.add_page_break()
H('PHẦN I — VIỆT NAM', level=1)
P('Hệ thống NC hiện có 58 bài, trong đó CHỈ 1 bài về CAN THIỆP tâm lý của tác giả Việt Nam — và đối tượng KHÔNG PHẢI VTN. Đây là khoảng trống nghiên cứu cực lớn.', bold=True, italic=True)

H('1.1. Trần Nguyễn Ngọc (2018) — Liệu pháp Thư giãn–Luyện tập cho RLLA lan tỏa', level=2)
H('Thông tin thư mục', level=3)
table(
    ['Mục', 'Chi tiết'],
    [
        ['Tác giả', 'Trần Nguyễn Ngọc'],
        ['Hướng dẫn KH', 'PGS.TS. Nguyễn Kim Việt'],
        ['Loại tài liệu', 'Luận án Tiến sĩ Y học, chuyên ngành Tâm thần (Mã số 62720148)'],
        ['Đơn vị', 'Trường Đại học Y Hà Nội — Viện Sức khỏe Tâm thần, BV Bạch Mai'],
        ['Năm', '2018, 177 trang'],
        ['Đối tượng', '170 bệnh nhân RLLA lan tỏa (F41.1 ICD-10), nội trú Viện SKTT Bạch Mai'],
        ['Đặc điểm mẫu', '65 nam (38,2%) + 105 nữ (61,8%) — NGƯỜI LỚN, KHÔNG PHẢI VTN'],
        ['Can thiệp', 'Liệu pháp Thư giãn–Luyện tập 20 buổi / 4 tuần (60 phút/buổi, 5-7 BN/nhóm)'],
        ['Thiết kế', 'Mục tiêu 1: cắt ngang mô tả lâm sàng (n=170) | Mục tiêu 2: can thiệp lâm sàng trước-sau (n=99)'],
        ['Nguyên tắc', 'KHÔNG kết hợp thuốc — đánh giá hiệu quả thuần của liệu pháp tâm lý-thể chất'],
        ['Công cụ đánh giá', 'HAM-A (Hamilton Anxiety) + CGI + PSQI (giấc ngủ) + EPI (nhân cách Eysenck)'],
    ],
    widths=[3.5, 12.0]
)

H('Mô tả can thiệp', level=3)
P('Liệu pháp Thư giãn–Luyện tập gồm 3 thành phần chính:', bold=True)
P('(1) Phần TẬP THƯ GIÃN: kỹ thuật thư giãn cơ tiến triển (Jacobson), thư giãn tự sinh (Schultz).')
P('(2) Phần TẬP THỞ: kỹ thuật thở bụng, thở chậm sâu, thở luân phiên (theo nguyên tắc Yoga).')
P('(3) Phần TẬP YOGA: các tư thế (asana) cơ bản phù hợp BN tâm thần — không vận động cường độ cao.')
P('Cấu trúc 1 buổi 60 phút chia 5 phần: (a) 15p đánh giá+thảo luận; (b) 10p khởi động; (c) 20p tập chính (thư giãn+thở+Yoga); (d) 10p tập trung tâm trí; (e) 5p kết thúc, đo sinh tồn. Có 1 bác sĩ + 1 cán bộ tâm lý + 1 điều dưỡng/buổi. Thực hiện tại Phòng Tâm lý lâm sàng Viện SKTT Bạch Mai.', italic=True)

H('Kết quả chính (n=99 hoàn thành 20 buổi)', level=3)
table(
    ['Chỉ số', 'T0 (vào viện)', 'T2 (tuần 2)', 'T4 (tuần 4)', 'p (T0-T4)'],
    [
        ['HAM-A — mức nặng', '45,5%', '22,2%', '11,1%', '<0,0001'],
        ['HAM-A — mức vừa', '20,2%', '24,2%', '36,4%', '<0,0001'],
        ['HAM-A — mức nhẹ', '34,3%', '53,6%', '52,5%', '0,0001'],
        ['Tổng số triệu chứng (X̄±SD)', '11,8 ± 3,5', '9,5 ± 3,8', '5,1 ± 4,9', '<0,0001'],
        ['Hồi hộp/tim đập nhanh', '88,9%', '73,7%', '43,4%', '<0,0001'],
        ['Vã mồ hôi', '59,6%', '36,3%', '16,1%', '<0,0001'],
        ['Run', '57,6%', '34,3%', '17,1%', '<0,0001'],
        ['Khô miệng', '38,4%', '25,2%', '16,2%', '<0,0001'],
        ['Khó thở', '56,6%', '46,4%', '25,2%', '<0,0001'],
    ],
    widths=[5.0, 2.5, 2.5, 2.5, 2.5]
)

H('Đặc điểm lâm sàng đáng chú ý (n=170)', level=3)
P('• 45,3% bệnh nhân có sang chấn tâm lý (32,4% sang chấn trường diễn) — chủ yếu sang chấn gia đình (33,5%), tai nạn/bệnh tật (27,1%).')
P('• 12,4% có TRẦM CẢM kèm theo; nữ cao hơn nam (15,2% vs 7,7%).')
P('• Sử dụng chất kết hợp: nam dùng thuốc lá (67,7%) và rượu (44,6%) cao hơn rõ rệt nữ (1,9% cả 2).')
P('• Theo Montgomery 2010 (cited): chỉ ~13,3% BN đến khám VÌ lo âu — phần lớn đến vì các triệu chứng cơ thể, dẫn đến chẩn đoán muộn.')

H('Quan điểm phản biện', level=3)
P('Điểm mạnh:', bold=True, color=RGBColor(0xCC,0,0))
P('• LUẬN ÁN ĐẦU TIÊN tại VN đánh giá hiệu quả liệu pháp Thư giãn–Luyện tập cho RLLA lan tỏa (cam đoan của tác giả: "Công trình này không trùng lặp với bất kỳ nghiên cứu nào khác đã được công bố tại Việt Nam").', color=RGBColor(0xCC,0,0))
P('• Mẫu LỚN n=170 (mô tả) + n=99 (can thiệp) — đủ cho thiết kế trước-sau.', color=RGBColor(0xCC,0,0))
P('• Quy trình chuẩn 20 buổi có giám sát bởi đa chuyên ngành (BS + tâm lý + điều dưỡng).', color=RGBColor(0xCC,0,0))
P('• Hiệu quả CÓ Ý NGHĨA THỐNG KÊ MẠNH (p<0,0001) trên CẢ HAM-A tổng và 22 triệu chứng riêng lẻ.', color=RGBColor(0xCC,0,0))
P('• Loại trừ thuốc — đánh giá thuần liệu pháp tâm lý-thể chất, có thể nhân rộng ở cơ sở y tế VN không có chuyên khoa tâm thần đầy đủ.', color=RGBColor(0xCC,0,0))
P('• Không tốn kém kinh tế, đơn giản, dễ thực hiện — phù hợp triển khai cộng đồng.', color=RGBColor(0xCC,0,0))

P('Hạn chế nghiêm trọng cho đề tài VTN:', bold=True, color=RGBColor(0xCC,0,0))
P('• ĐỐI TƯỢNG NGƯỜI LỚN — không phải VTN. KHÔNG THỂ ngoại suy trực tiếp cho HS THCS/THPT.', color=RGBColor(0xCC,0,0))
P('• THIẾT KẾ TRƯỚC-SAU không đối chứng — không phải RCT. Có thể có hiệu ứng giả dược, hồi quy về trung bình, và hiệu ứng Hawthorne (tham gia NC).', color=RGBColor(0xCC,0,0))
P('• Mẫu nội trú Bạch Mai — chọn lọc nghiêm trọng (BN đến viện thường nặng, có động lực cao), không đại diện cộng đồng.', color=RGBColor(0xCC,0,0))
P('• Theo dõi chỉ 4 tuần — không đo hiệu quả dài hạn (3 tháng, 6 tháng, 1 năm). So với Zhang 2026 Bayesian (đo 1 năm) — yếu hơn về thời gian.', color=RGBColor(0xCC,0,0))
P('• Không có nhóm đối chứng (chờ đợi/thuốc/CBT) — không thể so sánh hiệu lực với các phương pháp khác.', color=RGBColor(0xCC,0,0))
P('• 71/170 BN không hoàn thành nghiên cứu can thiệp (170 → 99) — tỷ lệ rút khỏi 41,8% RẤT CAO. Có thể có thiên lệch người sống sót (survivor bias).', color=RGBColor(0xCC,0,0))

H('Áp dụng cho đề tài VTN VN', level=3)
P('• Liệu pháp Thư giãn–Luyện tập có thể là MODULE BỔ TRỢ cho can thiệp CBT trường (giai đoạn 2 đề cương) — phù hợp với BMC NMA 2025 (PE SUCRA 0,51) và Cao 2025 Resilience MA.', bold=True)
P('• Nhưng cần ADAPTATION cho VTN: thời gian buổi ngắn hơn (30-45p), nội dung trẻ hơn (game/nhạc), kết hợp peer learning. Cần RCT có đối chứng.')
P('• KHOẢNG TRỐNG SỐ 1: Cần luận án/RCT VN ĐẦU TIÊN về can thiệp tâm lý cho VTN có RLLA — đối tượng HS THCS/THPT chưa được nghiên cứu can thiệp tại VN. Đây chính là trọng tâm đề cương của chúng ta.')

# ==========================================================================
# II. CHÂU Á (ngoài VN)
# ==========================================================================
doc.add_page_break()
H('PHẦN II — CHÂU Á (ngoài Việt Nam)', level=1)
P('7 công trình từ Trung Quốc, Nhật Bản, Sri Lanka, Indonesia/ĐNA, và scoping LMIC khu vực Đông Á-Thái Bình Dương. Đây là vùng có tính TƯƠNG ĐỒNG văn hóa cao nhất với VN, nên có giá trị tham chiếu trực tiếp.', italic=True)

H('Bảng II.1. Tổng quan 7 NC can thiệp Châu Á', level=2)
table(
    ['#', 'Tác giả/Năm', 'Quốc gia', 'Loại NC', 'Đối tượng', 'Phát hiện chính'],
    [
        ['41', 'Praptomojati 2024 (B7)', 'ĐNA tổng hợp', 'SR 7 NC', 'Trẻ + VTN', 'CA-CBT (CBT thích ứng văn hóa) hiệu quả; cần thích ứng tôn giáo + ngôn ngữ'],
        ['42', 'De Silva 2024 (B8)', 'Sri Lanka', 'Cluster RCT', '720 HS trường', 'CBT do GIÁO VIÊN cung cấp tại trường — lo âu giảm sau 3 tháng; mất mẫu <1%'],
        ['43', 'Xian 2024 (B9)', 'TQ-led', 'NMA 30 RCT', '1.547 VTN SAD', 'iCBT SUCRA 71,2% HẠNG 1; gCBT 68,4% hạng 2 (chức năng tốt nhất 89,6%)'],
        ['51', 'Japan iCBT 2024 (B11)', 'Nhật Bản', 'RCT đa trung tâm', 'VTN subthreshold SAD', 'iCBT cho SAD nhẹ — DƯƠNG TÍNH; mô hình prevention sớm'],
        ['57*', 'Qiaochu 2025 (B4)', 'TQ', 'SR 9 RCT', '2.479 trẻ + VTN', 'Mobile CBT: TRẦM CẢM 7/8 NC (87,5%) dương tính; LO ÂU 2/6 (33%) — yếu cho lo âu'],
        ['58*', 'Menon 2025 (B10)', 'LMIC ĐÁ-TBD', 'Scoping 69 NC', 'Trẻ + VTN, 12 nước', 'Tập trung "individual capacity" + clinical mgmt; THIẾU cộng đồng + gia đình + dài hạn'],
        ['29', 'Li 2025 BMC NMA (QT29)', 'TQ-led', 'NMA can thiệp lo âu', 'Trẻ', 'CBT SUCRA 0,66 (hạng 2); PE SUCRA 0,51; kết hợp đa modality tốt nhất'],
    ],
    widths=[1.0, 3.5, 2.5, 2.5, 3.5, 7.5]
)
P('* Bài 57, 58: ABSTRACT-ONLY (paywall — chưa có full PDF)', italic=True, size=10)

H('II.1. Praptomojati 2024 (#41) — CA-CBT cho ĐNA', level=2)
P('SR 7 nghiên cứu về CBT thích ứng văn hóa (Culturally Adapted CBT) tại Đông Nam Á (Indonesia, Malaysia, Philippines, Thái Lan...). Phát hiện: cần thích ứng (a) NGÔN NGỮ địa phương; (b) yếu tố TÔN GIÁO (Hồi giáo, Phật giáo, Kitô giáo); (c) HÌNH ẢNH metaphor phù hợp văn hóa; (d) vai trò GIA ĐÌNH mở rộng.')
P('Phản biện: SR khu vực ĐNA RẤT QUÝ cho VN, nhưng 7 NC vẫn là số lượng nhỏ; chưa có RCT chất lượng cao từ VN trong số này — VN là KHOẢNG TRỐNG ngay trong khu vực ĐNA. Áp dụng VN: cần phát triển CA-CBT phiên bản tiếng Việt + thích ứng văn hóa Á Đông + tôn giáo (Phật giáo, Công giáo).', italic=True, color=RGBColor(0xCC,0,0))

H('II.2. De Silva 2024 (#42) — Sri Lanka cluster RCT 720 HS', level=2)
P('Thiết kế CLUSTER RCT — phân ngẫu nhiên TRƯỜNG học (không phải HS), 720 HS THCS Sri Lanka. Can thiệp: CBT do GIÁO VIÊN cung cấp (sau khi được đào tạo). Theo dõi 3 tháng. Lo âu GIẢM CÓ Ý NGHĨA. Mất mẫu <1% (bằng chứng mạnh về tính khả thi).')
P('Quan trọng: Sri Lanka là LMIC Nam Á — bối cảnh nguồn lực hạn chế tương tự VN. Mô hình "GV cung cấp CBT" rẻ tiền, có thể nhân rộng. Tự trọng tăng β=0,811 sau can thiệp.')
P('Phản biện: Khác B5 UK (chuyên gia tâm lý > GV) — nhưng Sri Lanka cho thấy GV vẫn ĐỦ với đào tạo phù hợp. Đây là MÔ HÌNH KHẢ THI nhất cho VN. Cần lặp lại tại VN.', italic=True, color=RGBColor(0xCC,0,0))

H('II.3. Xian 2024 (#43) — Network Meta-Analysis SAD ở VTN', level=2)
P('NMA 30 RCT, n = 1.547 VTN có Lo âu Xã hội (SAD). So sánh đồng thời nhiều can thiệp:', bold=True)
table(
    ['Hạng', 'Can thiệp', 'SUCRA — triệu chứng', 'SUCRA — chức năng', 'Ghi chú'],
    [
        ['1', 'iCBT (internet CBT)', '71,2%', '—', 'Dẫn đầu — phù hợp digital-native VTN'],
        ['2', 'gCBT (CBT nhóm)', '68,4%', '89,6%', 'Tốt nhất cho phục hồi chức năng XH'],
        ['3', 'iCBT (cho lứa nhỏ)', '~60%', '—', 'Hiệu quả cả với trẻ <13 tuổi'],
        ['4', 'CBT cá nhân', '~50%', '—', 'Tiêu chuẩn vàng nhưng tốn kém'],
        ['5', 'Liệu pháp khác', '<40%', '—', 'Yếu hơn CBT'],
    ],
    widths=[1.5, 4.0, 3.5, 3.5, 4.0]
)
P('Phản biện: NMA cho phép so sánh GIÁN TIẾP tất cả can thiệp — bằng chứng MẠNH NHẤT về xếp hạng. iCBT đứng số 1 cho SAD — nên ưu tiên cho VN. gCBT số 2 nhưng RẤT TỐT cho chức năng XH. Hạn chế: chỉ SAD (không phải GAD), 30 RCT chủ yếu phương Tây + Đông Á phát triển.', italic=True, color=RGBColor(0xCC,0,0))

H('II.4. Japan iCBT 2024 (#51) — RCT đa trung tâm subthreshold SAD', level=2)
P('RCT đa trung tâm tại Nhật Bản. Đối tượng: VTN có triệu chứng SAD DƯỚI NGƯỠNG chẩn đoán (subthreshold). Can thiệp: iCBT 8-12 tuần, có hướng dẫn (guided). Kết quả DƯƠNG TÍNH — giảm triệu chứng, ngăn chuyển sang SAD đầy đủ.')
P('Quan trọng: Mô hình PREVENTION SỚM — can thiệp trước khi đủ tiêu chuẩn chẩn đoán. Phù hợp với chiến lược "stepped care" và "indicated prevention". Phù hợp văn hóa Á Đông (giảm kỳ thị nhờ hình thức online).')
P('Phản biện: RCT đa trung tâm — bằng chứng cao. Subthreshold là dân số quan trọng (rất lớn). VN có thể áp dụng tương tự với app iCBT tiếng Việt.', italic=True, color=RGBColor(0xCC,0,0))

H('II.5. Qiaochu 2025 (#57, abstract-only) — Mobile CBT 9 RCT TQ', level=2)
P('SR 9 RCT, N=2.479 trẻ + VTN. Can thiệp Mobile CBT (qua smartphone app).')
P('PHÁT HIỆN QUAN TRỌNG: 7/8 NC đo trầm cảm cho thấy giảm có ý nghĩa (87,5%). NHƯNG chỉ 2/6 NC đo lo âu cho thấy hiệu quả (33%). → Mobile CBT MẠNH cho TRẦM CẢM, YẾU cho LO ÂU (khi đo theo mô hình chung).')
P('Phản biện: Phát hiện này mâu thuẫn với B11 Japan iCBT (dương tính SAD) và B2 Walder JMIR DMHI MA (g=0,878 cho SAD-specific). Lý do có thể: (a) các app trong Qiaochu là cho TRẦM CẢM tổng hợp, không SAD-specific; (b) mobile vs internet (smartphone app vs web-based khác nhau). Cần đối chiếu thêm full PDF khi có.', italic=True, color=RGBColor(0xCC,0,0))

H('II.6. Menon 2025 (#58, abstract-only) — Scoping LMIC Đông Á-Thái Bình Dương', level=2)
P('Scoping Review 69 NC từ 12 quốc gia LMIC Đông Á + Thái Bình Dương: 32 RCT, 31 NC trước-sau, 6 đánh giá sau can thiệp. Bao gồm Việt Nam.')
P('Phát hiện chính: Có sự TẬP TRUNG MẤT CÂN ĐỐI vào "individual capacity" (năng lực cá nhân) và "clinical management". KHOẢNG TRỐNG ở: (a) thúc đẩy SK dựa CỘNG ĐỒNG; (b) phòng ngừa cấp GIA ĐÌNH; (c) dịch vụ ĐÁP ỨNG dài hạn. Phần lớn NC từ TQ + ĐNA; các nước nhỏ hơn + Thái Bình Dương đại diện tối thiểu.')
P('Phản biện: Cảnh báo MẠNH cho đề tài VN — không nên chỉ dừng ở can thiệp cá nhân (CBT). Cần thiết kế can thiệp ĐA CẤP: cá nhân + gia đình + cộng đồng + dài hạn. Phù hợp với phát hiện Dong 2025 (kênh giao tiếp gia đình OR=0,22) và Chen 2025 (family support là buffer chính).', italic=True, color=RGBColor(0xCC,0,0))

H('II.7. Li 2025 BMC NMA (#29) — Network Meta-Analysis can thiệp lo âu trẻ', level=2)
P('Network MA so sánh nhiều phương pháp can thiệp cho lo âu ở trẻ em. Kết quả SUCRA:')
P('• CBT: SUCRA 0,66 (hạng 2 — chỉ sau kết hợp đa modality)')
P('• PE (Hoạt động thể chất): SUCRA 0,51')
P('• Mindfulness: SUCRA ~0,40 (yếu)')
P('• Tâm lý động học: SUCRA <0,30')
P('• KẾT HỢP CBT + PE + family: hứa hẹn nhất')
P('Phản biện: Bằng chứng mạnh ủng hộ CBT là phương pháp tâm lý hiệu quả nhất. PE đứng độc lập cũng đáng kể — nên tích hợp vào can thiệp trường VN. Mindfulness yếu — phù hợp B5 UK 8.376 HS thất bại.', italic=True, color=RGBColor(0xCC,0,0))

# ==========================================================================
# III. ÂU - ÚC - MỸ
# ==========================================================================
doc.add_page_break()
H('PHẦN III — CHÂU ÂU - CHÂU ÚC - CHÂU MỸ', level=1)
P('7 công trình từ UK, Mỹ, Úc, châu Âu — bao gồm cả các MA quốc tế (chủ yếu RCTs phương Tây). Đây là vùng có bằng chứng lâu dài và phong phú nhất, cung cấp nền tảng phương pháp luận cho VN.', italic=True)

H('Bảng III.1. Tổng quan 7 NC can thiệp Âu-Úc-Mỹ', level=2)
table(
    ['#', 'Tác giả/Năm', 'Vùng', 'Loại NC', 'Mẫu', 'Phát hiện chính'],
    [
        ['44', 'Walder 2025 (B2)', 'Quốc tế', 'MA 21 RCT JMIR', 'VTN+người trẻ SAD', 'DMHI g=0,508; SAD-specific g=0,878; có hướng dẫn g=0,825 vs không 0,3'],
        ['49', 'JAMA App CBT 2024 (B3)', 'Mỹ', 'RCT', 'VTN', 'Mobile App CBT cho lo âu — RCT JAMA Open'],
        ['48', 'UK Brown & Carter 2025 (B5)', 'UK', 'Editorial Q1', '—', 'CBT > mindfulness; MHST + PLACES self-referral; mindfulness 8.376 HS THẤT BẠI; BESST 900 HS DƯƠNG TÍNH'],
        ['50', 'Cao 2025 Resilience MA (B6)', 'Quốc tế', 'SR + MA RCTs', 'Trẻ + VTN', 'Resilience trường effect nhỏ-TB; heterogeneity cao'],
        ['56*', 'Zhang 2026 Bayesian MA (B1)', 'Quốc tế', 'Bayesian MA 31 RCT', '19.865 HS', 'CBT phổ quát NHỎ; chất lượng nền THẤP; duy trì 1 năm'],
        ['28', 'Zugman 2024 AJP (QT28)', 'Mỹ', 'Review treatment', 'VTN', 'CAMS Walkup 2008 NEJM: CBT đơn 59,7%; CBT+SSRI 80,7%'],
        ['—', 'Walkup et al. 2008 NEJM (CAMS)', 'Mỹ', 'RCT đa trung tâm', '488 trẻ 7-17t', 'GOLD STANDARD: CBT+SSRI 80,7% > CBT 59,7% > SSRI 54,9% > placebo 23,7%'],
    ],
    widths=[1.0, 4.0, 2.0, 3.0, 2.5, 7.5]
)
P('* Bài 56: ABSTRACT-ONLY (paywall)', italic=True, size=10)

H('III.1. Walder 2025 JMIR (#44) — Digital Mental Health Interventions cho SAD', level=2)
P('Meta-analysis 21 RCT về DMHI (Digital Mental Health Interventions — internet, app, VR) cho Lo âu Xã hội (SAD) ở VTN + người trẻ. Tổng cộng nhiều nghìn người tham gia.')
P('Hệ số tác động (Hedges g):', bold=True)
P('• DMHI tổng quát: g = 0,508 (effect size TRUNG BÌNH)')
P('• DMHI SAD-specific (thiết kế riêng cho SAD): g = 0,878 (effect size LỚN)')
P('• DMHI có HƯỚNG DẪN người: g = 0,825 (effect size LỚN)')
P('• DMHI KHÔNG hướng dẫn (tự học hoàn toàn): g ≈ 0,3 (effect size NHỎ)')
P('Kết luận: DMHI hiệu quả tốt cho SAD VTN, NHƯNG phải (1) thiết kế ĐẶC THÙ cho SAD và (2) có hướng dẫn người (guided), KHÔNG chỉ là app tự học. Hai điều kiện này tăng hiệu lực gấp ~3 lần.')
P('Phản biện: Bằng chứng MẠNH NHẤT cho DMHI/iCBT trong SAD. Phù hợp Xian 2024 NMA hạng 1 và Japan iCBT B11. CHO VN: nếu phát triển app tiếng Việt, cần (a) dành riêng cho SAD/lo âu, (b) có hỗ trợ người (chat hỗ trợ, video call định kỳ).', italic=True, color=RGBColor(0xCC,0,0))

H('III.2. JAMA Network Open 2024 (#49) — Mobile App CBT cho lo âu', level=2)
P('RCT công bố trên JAMA Network Open 2024. Đối tượng: VTN có triệu chứng lo âu. Can thiệp: ứng dụng smartphone CBT-based.')
P('Phản biện: RCT JAMA — bằng chứng chất lượng cao. Phù hợp xu hướng digital health. Cần đối chiếu với Qiaochu B4 (mobile CBT yếu cho lo âu — có thể app này thiết kế tốt hơn).', italic=True, color=RGBColor(0xCC,0,0))

H('III.3. Brown & Carter 2025 — Editorial UK (#48)', level=2)
P('Editorial Journal of Mental Health (Q1, IF~3,5). Tổng hợp các mô hình can thiệp tại trường UK + đề xuất hướng đi.')
table(
    ['Mô hình', 'Người cung cấp', 'Hiệu quả', 'Bằng chứng'],
    [
        ['Mindfulness phổ quát (universal)', 'Giáo viên', 'THẤT BẠI', '8.376 HS / 85 trường (Kuyken 2022) — engagement thấp'],
        ['BESST (CBT self-referral)', 'Chuyên gia tâm lý', 'DƯƠNG TÍNH', '900 HS / 57 trường — Brown 2024'],
        ['MHST (Mental Health Support Team)', 'Thạc sĩ trị liệu', 'HỨA HẸN', 'Mới triển khai — chuyên gia trường'],
        ['PLACES (self-referral)', 'Đa cấp', 'HỨA HẸN', 'Dùng từ thường ngày, giảm kỳ thị'],
    ],
    widths=[5.0, 4.0, 3.0, 4.0]
)
P('Bài học: (a) Universal MINDFULNESS thất bại do engagement thấp; (b) TARGETED CBT (HS có triệu chứng tự đăng ký) thành công; (c) Chuyên gia > GV nhưng cả 2 đều khả thi; (d) Co-design (HS tham gia thiết kế) tăng engagement; (e) Ngôn ngữ thường ngày ("căng thẳng" thay "trầm cảm") giảm kỳ thị.')
P('Phản biện: Editorial — không phải SR/MA gốc, ý kiến chuyên gia + tổng hợp bằng chứng. Tuy nhiên cung cấp khung phân loại rất hữu ích. Cảnh báo về universal interventions là QUAN TRỌNG cho thiết kế đề cương VN.', italic=True, color=RGBColor(0xCC,0,0))

H('III.4. Cao 2025 Resilience MA (#50) — SR+MA RCTs Frontiers', level=2)
P('Tổng quan hệ thống + meta-analysis các RCT về can thiệp RESILIENCE (khả năng phục hồi) tại trường cho trẻ em + VTN. Frontiers in Psychiatry Q1 IF~4,7.')
P('Phương pháp: Random-effects MA. Đánh giá Cochrane Risk of Bias.')
P('Kết quả: Can thiệp resilience CÓ HIỆU QUẢ tăng resilience + giảm triệu chứng SKTT. NHƯNG kích thước hiệu ứng NHỎ-TRUNG BÌNH và heterogeneity CAO.')
P('Resilience khác CBT: tăng yếu tố BẢO VỆ (lạc quan, kết nối, kỹ năng ứng phó) thay vì giảm triệu chứng — bổ sung cho nhau.')
P('Phản biện: Q1 SR+MA RCTs — bằng chứng tốt. Effect nhỏ-TB nhưng resilience là YẾU TỐ BẢO VỆ — phù hợp với VN21 Trần Thảo Vi (lạc quan trung gian) và Ireland QT32 (resilience + tự trọng quan trọng tăng theo thời gian). Gợi ý: tích hợp module resilience vào can thiệp CBT VN.', italic=True, color=RGBColor(0xCC,0,0))

H('III.5. Zhang 2026 Bayesian MA (#56, abstract-only) — Long-term CBT trường', level=2)
P('Bayesian MA 31 RCT, n = 19.865 trẻ + VTN nguy cơ thấp (low-risk). CBT phổ quát tại trường cho trầm cảm + lo âu. Theo dõi dài hạn.')
P('Kết quả: Cải thiện CÓ Ý NGHĨA THỐNG KÊ nhưng KHIÊM TỐN trên trầm cảm; giảm NHỎ trên lo âu. Hiệu quả DUY TRÌ tới 1 năm.')
P('Kết luận tác giả: "Chất lượng bằng chứng nền RẤT THẤP khiến các phát hiện CHƯA ĐỦ vững chắc để hỗ trợ triển khai rộng rãi tại thời điểm này."')
P('Phản biện: CẢNH BÁO QUAN TRỌNG — không nên triển khai universal CBT ngay tức khắc dù được nhiều khuyến nghị quốc tế. Phù hợp B5 UK (mindfulness universal thất bại) và B1 nhấn mạnh cần CBT TARGETED. Cho VN: nên thiết kế can thiệp TARGETED (HS có triệu chứng) thay vì PHỔ QUÁT toàn trường.', italic=True, color=RGBColor(0xCC,0,0))

H('III.6. Zugman 2024 AJP (#28) — Pediatric anxiety treatment review', level=2)
P('Review tổng quan đăng American Journal of Psychiatry — về phương pháp tiếp cận điều trị lo âu nhi khoa hiện tại và tương lai.')
P('Tổng hợp bằng chứng từ CAMS (Walkup 2008 NEJM):', bold=True)
table(
    ['Nhánh điều trị', 'Tỷ lệ đáp ứng', 'Ý nghĩa'],
    [
        ['CBT + SSRI (Sertraline)', '80,7%', 'TỐT NHẤT — kết hợp tâm lý + thuốc'],
        ['CBT đơn thuần', '59,7%', 'TỐT — phương pháp tâm lý duy nhất'],
        ['SSRI đơn thuần (Sertraline)', '54,9%', 'KHÁ — thuốc duy nhất'],
        ['Placebo', '23,7%', 'BASELINE'],
    ],
    widths=[5.0, 3.0, 7.0]
)
P('Phản biện: CAMS là RCT đa trung tâm 488 trẻ 7-17 tuổi — GOLD STANDARD. Mặc dù từ 2008 nhưng vẫn được trích dẫn rộng rãi. Hạn chế: chỉ Mỹ; có thể không generalizable cho LMIC (chi phí thuốc + chuyên gia). Cho VN: nên ưu tiên CBT đơn thuần (không SSRI) — phù hợp với năng lực chuyên môn + chi phí + đạo đức cho VTN.', italic=True, color=RGBColor(0xCC,0,0))

# ==========================================================================
# TỔNG KẾT VÀ KHUYẾN NGHỊ
# ==========================================================================
doc.add_page_break()
H('PHẦN IV — TỔNG KẾT & KHUYẾN NGHỊ CHO VIỆT NAM', level=1)

H('IV.1. So sánh 3 vùng', level=2)
table(
    ['Tiêu chí', 'Việt Nam', 'Châu Á (ngoài VN)', 'Âu - Úc - Mỹ'],
    [
        ['Số bài can thiệp', '1 (Trần Nguyễn Ngọc 2018)', '7', '7'],
        ['Đối tượng VTN', 'KHÔNG (người lớn)', 'CÓ (đa số)', 'CÓ (đa số)'],
        ['Loại NC chủ yếu', 'Trước-sau không đối chứng', 'RCT + NMA + cluster RCT', 'RCT + MA + Bayesian MA'],
        ['Phương pháp chính', 'Thư giãn-Luyện tập (Yoga + thở)', 'CBT + iCBT + Mobile CBT + CA-CBT', 'CBT + iCBT + DMHI + Mindfulness'],
        ['Bối cảnh', 'Bệnh viện (nội trú)', 'Trường + Bệnh viện + Online', 'Trường + Online + Phòng khám'],
        ['Bằng chứng cao nhất', 'Yếu (1 luận án)', 'Mạnh (NMA Xian 2024)', 'Rất mạnh (CAMS NEJM, Bayesian MA)'],
        ['Khoảng trống', 'CỰC LỚN — 0 RCT VTN', 'Vừa — thiếu cộng đồng + gia đình', 'Nhỏ — đã có nhiều bằng chứng'],
    ],
    widths=[3.0, 4.0, 4.5, 4.5]
)

H('IV.2. Xếp hạng can thiệp dựa trên tổng hợp bằng chứng', level=2)
table(
    ['Hạng', 'Can thiệp', 'Bằng chứng tốt nhất', 'Phù hợp VN', 'Khuyến nghị'],
    [
        ['1', 'CBT + SSRI kết hợp', 'CAMS NEJM 80,7%', 'Hạn chế (chi phí + đạo đức)', 'Chỉ ca nặng có chỉ định'],
        ['2', 'CBT cá nhân/nhóm', 'CAMS 59,7%; BMC NMA SUCRA 0,66', 'CÓ (cốt lõi)', 'CHÍNH — đầu tư đào tạo'],
        ['3', 'iCBT cho SAD', 'Xian NMA hạng 1; Walder g=0,878', 'CAO — VTN digital-native', 'Phát triển app tiếng Việt'],
        ['4', 'gCBT (CBT nhóm)', 'Xian NMA hạng 2 chức năng', 'CAO — chi phí thấp', 'Triển khai trường'],
        ['5', 'CBT do GV (school-based)', 'Sri Lanka B8 RCT 720 HS', 'CAO — LMIC khả thi', 'Đào tạo GV chuyên môn'],
        ['6', 'CA-CBT (thích ứng văn hóa)', 'Praptomojati ĐNA SR', 'CỰC CAO', 'Phát triển phiên bản VN'],
        ['7', 'Hoạt động thể chất (PE)', 'BMC NMA SUCRA 0,51', 'CAO — sẵn có ở trường', 'Tích hợp giờ thể dục'],
        ['8', 'Resilience training', 'Cao 2025 MA Frontiers', 'CAO', 'Module bổ trợ CBT'],
        ['9', 'Thư giãn-Luyện tập (Yoga)', 'Trần Nguyễn Ngọc VN 2018', 'CAO — đã thử nghiệm VN', 'Module bổ trợ; cần RCT VTN'],
        ['10', 'Mobile CBT (trầm cảm)', 'Qiaochu 7/8 NC dương tính', 'CAO', 'Cho TRẦM CẢM mạnh hơn lo âu'],
        ['11', 'PLACES self-referral', 'B5 UK BESST 900 HS', 'TB — cần adapt', 'Thử nghiệm pilot'],
        ['12', 'Mindfulness phổ quát', 'B5 UK 8.376 HS THẤT BẠI', 'KHÔNG nên triển khai universal', 'Tránh universal'],
        ['13', 'Universal CBT phổ quát', 'Zhang Bayesian — hiệu quả nhỏ', 'KHÔNG nên — TARGETED tốt hơn', 'Tránh — chuyển TARGETED'],
    ],
    widths=[1.0, 4.5, 4.0, 3.5, 4.0]
)

H('IV.3. Khoảng trống nghiên cứu can thiệp tại VN', level=2)
P('Sau khi tổng hợp 15 bài can thiệp, 5 khoảng trống lớn nhất cho VN là:', bold=True)
P('1. RCT CAN THIỆP TÂM LÝ ĐẦU TIÊN cho VTN có RLLA tại VN — chưa có bài nào. Trần Nguyễn Ngọc 2018 là người lớn nội trú.')
P('2. Phát triển app iCBT/Mobile CBT TIẾNG VIỆT — phù hợp xếp hạng số 3-4 (Xian, Walder, Japan).')
P('3. Đào tạo GIÁO VIÊN cung cấp CBT tại trường — mô hình Sri Lanka khả thi cho LMIC.')
P('4. CA-CBT (CBT thích ứng văn hóa Á Đông) phiên bản VN — yếu tố Phật giáo, gia đình mở rộng, "fear of letting down" (Dong 2025).')
P('5. CAN THIỆP ĐA CẤP: cá nhân + GIA ĐÌNH + cộng đồng — lấp khoảng trống Menon B10 + phát hiện kênh giao tiếp gia đình Dong A5 (OR=0,22).')

H('IV.4. Đề xuất thiết kế can thiệp cho đề cương VN (giai đoạn 2)', level=2)
P('Dựa trên tổng hợp 15 bài, đề xuất can thiệp 12 tuần TARGETED cho HS THCS/THPT có triệu chứng GAD-7 ≥8 hoặc DASS-Y lo âu ≥8:', bold=True)
P('THÀNH PHẦN 1 — CBT NHÓM TARGETED (8 buổi, 60p/buổi, 8-12 HS/nhóm): Dựa trên CAMS, BMC NMA, Sri Lanka. Do giáo viên/cố vấn tâm lý đã đào tạo cung cấp.')
P('THÀNH PHẦN 2 — MODULE GIAO TIẾP GIA ĐÌNH (4 buổi, có cha mẹ): KIẾN THỨC MỚI từ Dong 2025 (OR=0,22) + Menon B10 (gap GIA ĐÌNH). Kỹ năng "tâm sự" — listening, không phán xét.')
P('THÀNH PHẦN 3 — RESILIENCE (3 buổi): Dựa trên Cao 2025 MA. Lạc quan, kết nối XH, tự nhận thức.')
P('THÀNH PHẦN 4 — HOẠT ĐỘNG THỂ CHẤT + THỞ-THƯ GIÃN (8 buổi tích hợp giờ thể dục): Dựa trên BMC NMA PE SUCRA 0,51 + Trần Nguyễn Ngọc 2018 (cho phép VN).')
P('THÀNH PHẦN 5 — APP iCBT TIẾNG VIỆT (suốt 12 tuần, hỗ trợ qua chat): Phát triển dựa trên Walder DMHI g=0,878 (SAD-specific + có hướng dẫn) + Japan B11.')
P('CƠ CHẾ TIẾP CẬN: PLACES self-referral (UK B5) — HS tự đăng ký dùng từ thường ngày ("căng thẳng", "lo lắng").')

H('IV.5. Kết luận', level=2)
P('Việt Nam đang có khoảng trống RẤT LỚN về nghiên cứu can thiệp tâm lý cho VTN có RLLA — 0 RCT trong hệ thống 58 bài. Trần Nguyễn Ngọc 2018 là cơ sở duy nhất từ VN nhưng đối tượng người lớn nội trú. Bằng chứng quốc tế (đặc biệt từ Châu Á) cho thấy CBT (cá nhân, nhóm, internet) là phương pháp hiệu quả nhất; mô hình Sri Lanka (GV cung cấp) khả thi cho LMIC; iCBT phù hợp VTN digital-native; PHẢI tránh universal interventions vô hiệu lực; cần tích hợp can thiệp GIA ĐÌNH (phát hiện mới Dong 2025). Đề cương VN nên thiết kế can thiệp TARGETED 12 tuần đa thành phần — đây sẽ là RCT đầu tiên loại này tại VN.', bold=True, italic=True)

P('---', italic=True)
P('Tài liệu tham chiếu: 15 bài can thiệp (#28, #29, #41-44, #48-51, #56-58 + Trần Nguyễn Ngọc 2018) + Walkup et al. 2008 NEJM (CAMS, lịch sử). Xem chi tiết bản dịch trong thư mục 03_Ban-dich/.', italic=True, size=10)
P('Cập nhật: 10/04/2026 | Tổng hợp bởi nhóm NC dự án Lo âu | Phong cách Công Thị Hằng v5', italic=True, size=10)

doc.save(OUT)
print(f'Saved: {OUT}')
d2 = Document(OUT)
print(f'Total paragraphs: {len(d2.paragraphs)}, tables: {len(d2.tables)}')
