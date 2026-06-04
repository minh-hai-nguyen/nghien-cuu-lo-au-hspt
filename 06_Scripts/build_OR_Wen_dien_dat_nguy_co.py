"""Build doc: dien dat OR = 11,579 (Wen 2020) — 'nguy co' co dung khong?"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/OR_Wen_2020_dien_dat_nguy_co_vs_odds.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(12); r.italic = italic
    return p

# =========================================================
# TIÊU ĐỀ
# =========================================================
H('Bình luận diễn đạt OR = 11,579 (Wen et al. 2020)', level=1)
H("Áp lực học tập rất cao ↔ Lo âu nặng — 'nguy cơ' hay 'odds'?", level=2)

# =========================================================
# CÂU HỎI (xanh)
# =========================================================
H('Câu hỏi của thầy', level=2)
para_blue(
    'Diễn giải số liệu này: X. Wen và cs. (2020) — áp lực học tập rất cao là yếu tố nguy cơ '
    'mạnh nhất (OR = 11,579 cho lo âu nặng so với áp lực rất thấp), tức là những học sinh '
    'chịu áp lực học tập rất cao có nguy cơ bị rối loạn lo âu nặng cao gấp khoảng 11,5 lần '
    'so với học sinh chịu áp lực học tập rất thấp (OR = 11,579). Diễn đạt lại như vậy có '
    'đúng không em?'
)

# =========================================================
# BACKGROUND (đen)
# =========================================================
H('1. Hai khái niệm khác nhau: ODDS và RISK', level=2)

para_black(
    'Trước khi nói đúng/sai, cần phân biệt rõ hai khái niệm trong dịch tễ học mà người Việt '
    'thường dịch chung là "nguy cơ":'
)

bullet_black(
    'RISK (nguy cơ tuyệt đối) = số người bị bệnh / tổng số người trong nhóm. '
    'Ví dụ: 100 HS áp lực rất cao, 30 em bị lo âu nặng → risk = 30%.'
)
bullet_black(
    'ODDS (chênh) = số người bị bệnh / số người KHÔNG bị bệnh. '
    'Cùng ví dụ trên: odds = 30/70 = 0,43.'
)

para_black('Tỷ số tương ứng cũng khác nhau:')
bullet_black('RR (Risk Ratio = tỷ số nguy cơ) = risk nhóm phơi nhiễm / risk nhóm không phơi nhiễm.')
bullet_black('OR (Odds Ratio = tỷ số chênh) = odds nhóm phơi nhiễm / odds nhóm không phơi nhiễm.')

para_black('Quan hệ giữa OR và RR:', bold=True)
bullet_black('Khi outcome HIẾM (prevalence < 10%): OR ≈ RR → có thể nói "nguy cơ gấp X lần" tương đương với "odds gấp X lần".')
bullet_black('Khi outcome PHỔ BIẾN (prevalence > 10%): OR PHÓNG ĐẠI so với RR. Outcome càng phổ biến, OR càng "lệch" cao hơn RR thật.')

H('2. Áp dụng vào con số OR = 11,579 của Wen 2020', level=2)

para_black('Lo âu nặng ở HS THCS/THPT KHÔNG phải outcome hiếm:')
bullet_black('VNAMHS 2022 (VN002): tỷ lệ lo âu chung ở trẻ em VN ~ 18,6%; lo âu nặng (clinical) ~ 5–8%.')
bullet_black('VN016 Nguyễn Cao Minh (2012) chuẩn hóa RCADS: ~10–15% có triệu chứng lo âu đáng chú ý.')
bullet_black('Trung Quốc (Wen 2020 và các nghiên cứu cùng thời): lo âu nặng ở HS dao động 7–15% tùy thang đo.')

para_black(
    'Vì prevalence ~ 5–15% (vùng "không hiếm"), OR = 11,579 sẽ PHÓNG ĐẠI so với RR thật. '
    'Ước tính nhanh (theo công thức Zhang & Yu 1998 hoặc bảng quy đổi Davies et al. 1998): '
    'với prevalence ~ 10%, OR = 11,5 tương đương RR ~ 4–6, không phải 11,5.'
)

H('3. Vì sao diễn đạt "nguy cơ gấp 11,5 lần" là PHỔ BIẾN nhưng KHÔNG CHÍNH XÁC', level=2)

para_black(
    'Trong báo chí khoa học, sách giáo trình phổ thông, các bài giảng đại học (kể cả y học), '
    'người ta thường viết "nguy cơ tăng X lần" cho cả OR lẫn RR — đây là cách nói tắt '
    'CHẤP NHẬN ĐƯỢC khi:'
)
bullet_black('Người đọc là độc giả phổ thông, không cần phân biệt OR/RR chi tiết.')
bullet_black('Outcome là biến cố hiếm (OR ≈ RR).')
bullet_black('Có ghi chú rõ ở chỗ khác về OR và độ phóng đại.')

para_black('Tuy nhiên trong báo cáo NGHIÊN CỨU, LUẬN VĂN, ĐỀ TÀI KHOA HỌC nghiêm túc:', bold=True)
bullet_black('Phải dùng "odds" hoặc "khả năng tương đối" cho OR.')
bullet_black('Phải dùng "nguy cơ" chỉ cho RR (relative risk) hoặc absolute risk thật.')
bullet_black('Phải kèm KTC 95% và tỷ lệ outcome trong dân số gốc.')

H('4. Kiểm tra mức độ OR = 11,579', level=2)

tbl = doc.add_table(rows=6, cols=3)
tbl.style = 'Light Grid Accent 1'
header = tbl.rows[0]
header.cells[0].text = 'OR (hoặc 1/OR)'
header.cells[1].text = 'Mức độ'
header.cells[2].text = 'Vị trí của OR = 11,579'
data = [
    ('1,0 – 1,5', 'Rất yếu', ''),
    ('1,5 – 2,5', 'Yếu', ''),
    ('2,5 – 4,0', 'Trung bình', ''),
    ('4,0 – 6,0', 'Mạnh', ''),
    ('> 6,0', 'Rất mạnh / cực mạnh', '✓ OR = 11,579 thuộc nhóm này — VƯỢT XA ngưỡng "very strong"'),
]
for i, (a, b, c) in enumerate(data, 1):
    tbl.rows[i].cells[0].text = a
    tbl.rows[i].cells[1].text = b
    tbl.rows[i].cells[2].text = c
for row in tbl.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black(
    'OR = 11,579 là CỰC LỚN. Trong bối cảnh tâm lý vị thành niên, OR > 10 ở một yếu tố '
    'nguy cơ duy nhất là phát hiện đáng nghi ngờ về phương pháp — cần kiểm tra:',
    bold=True
)
bullet_black('Cỡ mẫu có đủ lớn không? OR cao đôi khi do cỡ mẫu nhỏ + ít người trong ô tham chiếu.')
bullet_black('KTC 95% rộng cỡ nào? Nếu KTC kiểu 4–25 thì điểm ước lượng 11,5 không đáng tin cậy.')
bullet_black('Đã hiệu chỉnh confounders chưa (đặc biệt: trầm cảm đồng mắc, giới, lớp, hoàn cảnh kinh tế)?')
bullet_black('Định nghĩa "áp lực rất cao" và "áp lực rất thấp" có khoảng cách thực sự lớn không, hay chỉ là 2 đầu phân cực của thang Likert?')

# =========================================================
# CÂU TRẢ LỜI (xanh, gom 1 chỗ trước phụ lục)
# =========================================================
H('5. CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Trả lời ngắn cho thầy:', bold=True)
para_blue(
    'Diễn đạt của thầy đúng về Ý NGHĨA TỔNG THỂ — học sinh áp lực học tập rất cao có khả '
    'năng (theo odds) bị lo âu nặng cao hơn nhiều so với áp lực rất thấp. Nhưng có MỘT '
    'CHỖ CẦN CHỈNH cho chính xác về thống kê: chữ "nguy cơ".'
)

para_blue('Vấn đề:', bold=True)
bullet_blue(
    'OR (Odds Ratio) ≠ RR (Risk Ratio). OR là tỷ số CHÊNH (odds), không phải tỷ số NGUY CƠ. '
    'Khi outcome không hiếm (lo âu nặng ở HS ~ 5–15%), OR phóng đại so với RR thật.'
)
bullet_blue(
    'Với OR = 11,5 và prevalence ~ 10%, RR thực tế ước tính chỉ khoảng 4–6 — tức nguy cơ '
    'thật (theo risk) gấp khoảng 4–6 lần, KHÔNG PHẢI 11,5 lần.'
)

para_blue('Hai cách diễn đạt khuyến nghị:', bold=True)

para_blue(
    'Cách 1 — chính xác thống kê (khuyên dùng cho báo cáo nghiên cứu, đề cương, luận văn):'
)
bullet_blue(
    '"Học sinh chịu áp lực học tập rất cao có ODDS (khả năng tương đối) bị rối loạn lo '
    'âu nặng GẤP KHOẢNG 11,5 LẦN so với học sinh chịu áp lực học tập rất thấp '
    '(OR = 11,579; KTC 95% = …)."'
)

para_blue(
    'Cách 2 — phổ thông (chấp nhận được cho báo cáo dạng phổ biến, kèm chú thích):'
)
bullet_blue(
    '"Học sinh chịu áp lực học tập rất cao có nguy cơ bị rối loạn lo âu nặng cao gấp '
    'khoảng 11,5 lần so với áp lực rất thấp (OR = 11,579). Lưu ý: con số này là tỷ số '
    'chênh (odds ratio); vì lo âu nặng không phải biến cố hiếm, nguy cơ tuyệt đối có thể '
    'thấp hơn — RR ước tính khoảng 4–6 lần."'
)

para_blue('Bình luận thêm về độ mạnh:', bold=True)
bullet_blue(
    'OR = 11,579 thuộc nhóm CỰC MẠNH (vượt xa ngưỡng "very strong" > 6). Đây là OR cao '
    'khác thường — cần kiểm tra KTC 95%, cỡ mẫu, cách định nghĩa "áp lực rất cao" và '
    '"rất thấp", và các confounders đã được hiệu chỉnh chưa.'
)
bullet_blue(
    'OR cực cao đôi khi do mẫu so sánh nhỏ ở cực dưới (số HS "áp lực rất thấp" ít → '
    'tỷ lệ lo âu nặng ở nhóm này gần 0 → odds nhóm tham chiếu rất nhỏ → OR phình ra). '
    'Khi viết báo cáo, thầy nên nêu cả KTC 95% để người đọc đánh giá độ tin cậy.'
)
bullet_blue(
    'Phát hiện rằng áp lực học tập là yếu tố nguy cơ MẠNH NHẤT phù hợp với corpus dự án '
    '(QT067 Pascoe 2020, VN013 Tô Thị Hồng 2017, các nghiên cứu Trung Quốc cùng thời). '
    'Đáng đưa vào báo cáo can thiệp.'
)

para_blue('Tóm lại:', bold=True)
para_blue(
    'Ý của thầy ĐÚNG. Chỉ cần đổi "nguy cơ" → "odds/khả năng" trong câu chính, hoặc giữ '
    'từ "nguy cơ" nhưng thêm chú thích về việc OR phóng đại so với RR khi outcome không '
    'hiếm. Đồng thời đính kèm KTC 95% mỗi khi trích OR = 11,579.'
)

# =========================================================
# PHỤ LỤC — REFERENCES
# =========================================================
H('6. Phụ lục — Tài liệu tham khảo', level=2)

para_black(
    '1. Wen, X., et al. (2020). [Bài gốc thầy đang trích — cần bổ sung citation đầy đủ '
    'từ PDF: tên đầy đủ, tạp chí, DOI]. — Khi đưa OR = 11,579 vào báo cáo cuối, em đề '
    'nghị thầy điền KTC 95% từ bài gốc để trích dẫn đầy đủ.',
    italic=True
)

para_black(
    '2. Zhang, J., & Yu, K. F. (1998). What\'s the relative risk? A method of correcting '
    'the odds ratio in cohort studies of common outcomes. JAMA, 280(19), 1690–1691. '
    'doi:10.1001/jama.280.19.1690 — Công thức kinh điển để chuyển đổi OR → RR khi outcome '
    'phổ biến.',
    italic=True
)

para_black(
    '3. Davies, H. T. O., Crombie, I. K., & Tavakoli, M. (1998). When can odds ratios '
    'mislead? BMJ, 316(7136), 989–991. doi:10.1136/bmj.316.7136.989 — Bài kinh điển cảnh '
    'báo về việc dùng OR thay RR; có bảng quy đổi nhanh.',
    italic=True
)

para_black(
    '4. Schmidt, C. O., & Kohlmann, T. (2008). When to use the odds ratio or the relative '
    'risk? International Journal of Public Health, 53(3), 165–167. — Hướng dẫn chọn OR '
    'hay RR theo loại nghiên cứu.',
    italic=True
)

para_black(
    '5. Hosmer, D. W., Lemeshow, S., & Sturdivant, R. X. (2013). Applied Logistic '
    'Regression (3rd ed.). Hoboken, NJ: Wiley. — Sách tham khảo chuẩn về OR.',
    italic=True
)

para_black('6. Tham khảo trong corpus dự án (DB Lo-au):')
bullet_black(
    'QT067 Pascoe, Hetrick & Parker (2020) — narrative review về academic stress; phù hợp '
    'với phát hiện áp lực học tập là nguy cơ mạnh.', italic=True
)
bullet_black(
    'VN002 VNAMHS 2022 — prevalence lo âu trẻ VN, dùng để ước tính RR thật từ OR.',
    italic=True
)
bullet_black(
    'VN013 Tô Thị Hồng (2017) — DASS-21 trên HS THCS Hà Nội; có dữ liệu OR theo áp lực '
    'học tập + parenting.', italic=True
)
bullet_black(
    'VN016 Nguyễn Cao Minh (2012) — RCADS chuẩn hóa VN; baseline prevalence để compare.',
    italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
