# -*- coding: utf-8 -*-
"""Sinh doc: Tai lieu tham khao dot 2 cho bai Q1 (8 bai moi, da xac minh). 19/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-khgdvn', 'Tai_lieu_tham_khao_Q1_dot2_19052026.docx')
PAGE_W = 16.0
GREEN = RGBColor(0x1F, 0x70, 0x33)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    return doc
def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)
    return h
def P(doc, text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p
def field(doc, label, text, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(label)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r2 = p.add_run(text)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(12)
    if color is not None: r2.font.color.rgb = color
    return p
def bullet(doc, text, size=11):
    p = doc.add_paragraph(style='List Bullet'); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    return p
def add_hyperlink(paragraph, url, text, size=11):
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    h = OxmlElement('w:hyperlink'); h.set(qn('r:id'), r_id)
    nr = OxmlElement('w:r'); rPr = OxmlElement('w:rPr')
    rf = OxmlElement('w:rFonts'); rf.set(qn('w:ascii'), 'Times New Roman'); rf.set(qn('w:hAnsi'), 'Times New Roman')
    rPr.append(rf)
    c = OxmlElement('w:color'); c.set(qn('w:val'), '0563C1'); rPr.append(c)
    u = OxmlElement('w:u'); u.set(qn('w:val'), 'single'); rPr.append(u)
    sz = OxmlElement('w:sz'); sz.set(qn('w:val'), str(size*2)); rPr.append(sz)
    nr.append(rPr)
    t = OxmlElement('w:t'); t.text = text; nr.append(t)
    h.append(nr); paragraph._p.append(h)
def table(doc, headers, rows, widths, hdr=9, body=9):
    assert abs(sum(widths)-PAGE_W) < 0.1, sum(widths)
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = htext
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(hdr)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(body)
    return t

doc = make_doc()
ttl = doc.add_paragraph(); ttl.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = ttl.add_run('TÀI LIỆU THAM KHẢO ĐỢT 2 — CHUẨN BỊ VIẾT BÀI Q1\n'
                'Tám bài mới đã xác minh (3 Việt Nam + 5 quốc tế)')
r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = sub.add_run('Ngày lập: 19/05/2026 — Nối tiếp doc "Bài tham khảo đã xác minh – chuẩn bị bài Q1" (6 bài đợt 1)')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

P(doc, 'Tài liệu này bổ sung 8 bài tham khảo cho bài Q1 — hướng nghiên cứu "đường dẫn yếu '
       'tố nguy cơ – bảo vệ tới lo âu/cảm xúc tiêu cực ở học sinh trung học Việt Nam", có '
       'định hướng nâng cấp lên thiết kế dọc và mô hình phương trình cấu trúc (Structural '
       'Equation Modeling — SEM). Tám bài được đánh mã nối tiếp đợt 1: A4–A6 (Việt Nam), '
       'B4–B8 (quốc tế). Toàn bộ đã được mở trực tiếp trang nguồn (PubMed/PMC/nhà xuất bản) '
       'để xác minh tác giả, năm, tạp chí, số bài và DOI; 7/8 bài đã tải được bản PDF toàn văn.')

H(doc, '1. Bảng tổng hợp 8 bài', 1)
table(doc,
    ['Mã', 'Tác giả (năm)', 'Nước', 'Thiết kế', 'Truy cập mở', 'PDF'],
    [
        ['A4', 'Nguyen T. T. (2024)', 'Việt Nam', 'Cắt ngang', 'Có', 'Đã tải'],
        ['A5', 'Tran L. C. T. và cs. (2025)', 'Việt Nam', 'Cắt ngang', 'Có', 'Đã tải'],
        ['A6', 'Ho, Nguyen & Nguyen (2022)', 'Việt Nam', 'Cắt ngang (trung gian có điều tiết)', 'Có', 'Đã tải'],
        ['B4', 'Chen và cs. (2025)', 'Trung Quốc', 'Cắt ngang (trung gian chuỗi)', 'Có', 'Đã tải'],
        ['B5', 'Wang và cs. (2025)', 'Trung Quốc', 'Dọc 3 đợt (trung gian)', 'Có', 'Đã tải'],
        ['B6', 'Šutić và cs. (2025)', 'Hà Lan', 'Dọc cường độ cao (DSEM)', 'Có', 'Đã tải'],
        ['B7', 'Yang và cs. (2024)', 'Trung Quốc', 'Cắt ngang (SEM đa trung gian)', 'Có', 'Đã tải'],
        ['B8', 'Tan và cs. (2025)', 'Trung Quốc', 'Cắt ngang (SEM)', 'Chưa rõ', 'Chưa tải'],
    ],
    widths=[1.0, 4.0, 1.9, 4.6, 1.9, 2.6])

# ================= A. VIET NAM =================
H(doc, '2. Ba bài Việt Nam', 1)

H(doc, 'A4. Nguyen (2024) — hỗ trợ xã hội là yếu tố bảo vệ, mẫu THCS', 2)
field(doc, 'Trích dẫn: ', 'Nguyen, T. T. (2024). Emotional-behavioral difficulties and '
      'prosocial behaviour among Vietnamese adolescents: The role of social support. '
      'Health Psychology Research, 12, Article 116967.')
field(doc, 'DOI / Nguồn: ', '10.52965/001c.116967 — PMC11093747.')
field(doc, 'Thiết kế & mẫu: ', 'Cắt ngang định lượng; 582 vị thành niên (287 nam, 295 nữ) '
      '12–15 tuổi (lớp 6–9), vùng Đồng bằng sông Cửu Long. Công cụ: Bảng hỏi Điểm mạnh và '
      'Khó khăn (Strengths and Difficulties Questionnaire — SDQ) đo khó khăn cảm xúc – hành '
      'vi và hành vi vì xã hội; thang đo hỗ trợ xã hội từ gia đình và bạn bè.')
field(doc, 'Phát hiện chính: ', 'Hỗ trợ từ gia đình dự báo giảm có ý nghĩa cả triệu chứng '
      'cảm xúc lẫn triệu chứng hành vi; hỗ trợ từ bạn bè chỉ bảo vệ khỏi vấn đề hành vi. Mọi '
      'nguồn hỗ trợ đều dự báo tăng hành vi vì xã hội. Hỗ trợ gia đình nổi lên là yếu tố bảo '
      'vệ chính.')
field(doc, 'Truy cập mở: ', 'Có — toàn văn miễn phí trên PMC.')
field(doc, '→ Giá trị cho bài Q1: ', 'Một trong số ít nghiên cứu Việt Nam đúng nhóm tuổi '
      'học sinh trung học cơ sở (12–15, lớp 6–9), với hỗ trợ xã hội là yếu tố bảo vệ — cung '
      'cấp dữ liệu nền cho tầng "gia đình/bạn bè" trong khung đường dẫn nguy cơ – bảo vệ.',
      color=GREEN)

H(doc, 'A5. Tran và cộng sự (2025) — yếu tố nguy cơ/bảo vệ, mẫu lớn, mới nhất', 2)
field(doc, 'Trích dẫn: ', 'Tran, L. C. T., Rattanapan, C., Laosee, O., Suksaroj, T. T., '
      'Du, R., Ramazanu, S., Truong, T. N., To, T. B. S., & Shorey, S. (2025). Prevalence '
      'of depressive symptoms and their determinants among adolescents in Can Tho City, '
      'Vietnam: A cross-sectional study. Cambridge Prisms: Global Mental Health, 12, '
      'Article e139.')
field(doc, 'DOI / Nguồn: ', '10.1017/gmh.2025.10096 — PMC12720387.')
field(doc, 'Thiết kế & mẫu: ', 'Cắt ngang; 1.054 vị thành niên 15–18 tuổi tại Thành phố Cần '
      'Thơ (Đồng bằng sông Cửu Long).')
field(doc, 'Phát hiện chính: ', '37,4% có triệu chứng trầm cảm. Yếu tố nguy cơ: méo mó '
      'nhận thức, áp lực học tập, bạo lực thể chất trong gia đình, người nhà từng bị giam '
      'giữ, sử dụng rượu và thuốc lá, bắt nạt học đường, bạo lực tình dục qua công nghệ. Yếu '
      'tố bảo vệ: không có tiền sử trầm cảm trong gia đình, không dùng rượu/thuốc lá, không '
      'bị bắt nạt, không bị bạo lực tình dục.')
field(doc, 'Truy cập mở: ', 'Có — toàn văn miễn phí.')
field(doc, '→ Giá trị cho bài Q1: ', 'Bài Việt Nam mới nhất (2025), mẫu lớn, phân tách rõ '
      'yếu tố nguy cơ và bảo vệ ở nhiều tầng — cập nhật dữ liệu nền; đăng cùng tạp chí '
      'Cambridge Prisms: Global Mental Health (uy tín, phù hợp định hướng Q1).', color=GREEN)

H(doc, 'A6. Ho, Nguyen & Nguyen (2022) — mô hình trung gian có điều tiết', 2)
field(doc, 'Trích dẫn: ', 'Ho, T. T., Nguyen, B. T. N., & Nguyen, N. P. H. (2022). Academic '
      'stress and depression among Vietnamese adolescents: A moderated mediation model of '
      'life satisfaction and resilience. Current Psychology. (Xuất bản trực tuyến trước, '
      '2022).')
field(doc, 'DOI / Nguồn: ', '10.1007/s12144-022-03661-3 — PMID 36277264 — PMC9574843.')
field(doc, 'Thiết kế & mẫu: ', 'Cắt ngang; 1.336 vị thành niên Việt Nam. Công cụ: Thang '
      'Stress Học tập cho Vị thành niên (ESSA), Thang Kiên cường Connor–Davidson (CD-RISC), '
      'Thang Hài lòng Cuộc sống (SWLS), Thang Trầm cảm Beck phiên bản 2 (BDI-II).')
field(doc, 'Phát hiện chính: ', 'Mô hình trung gian có điều tiết (moderated mediation): '
      'kiên cường (resilience) trung gian một phần quan hệ stress học tập → trầm cảm; hài '
      'lòng cuộc sống điều tiết có ý nghĩa hiệu ứng gián tiếp này.')
field(doc, 'Truy cập mở: ', 'Có — toàn văn miễn phí trên PMC.')
field(doc, '→ Giá trị cho bài Q1: ', 'Một trong ít bài Việt Nam dùng mô hình TRUNG GIAN CÓ '
      'ĐIỀU TIẾT với yếu tố bảo vệ (kiên cường, hài lòng cuộc sống) — khuôn mẫu phân tích '
      'trực tiếp cho khung đường dẫn của bài Q1. Lưu ý: kết cục đo là trầm cảm (BDI-II), '
      'không phải lo âu.', color=GREEN)

# ================= B. QUOC TE =================
H(doc, '3. Năm bài quốc tế', 1)

H(doc, 'B4. Chen và cộng sự (2025) — SEM trung gian chuỗi, bầu không khí trường học', 2)
field(doc, 'Trích dẫn: ', 'Chen, W., Huang, Z., Peng, B., & Hu, H. (2025). Unpacking the '
      'relationship between adolescents’ perceived school climate and negative '
      'emotions: The chain mediating roles of school belonging and social avoidance and '
      'distress. BMC Psychology, 13, Article 58.')
field(doc, 'DOI / Nguồn: ', '10.1186/s40359-025-02364-1 — PMID 39833960 — PMC11748506.')
field(doc, 'Thiết kế & mẫu: ', 'Cắt ngang; 1.507 vị thành niên Trung Quốc lớp 5–9 '
      '(khoảng 10–15 tuổi).')
field(doc, 'Phát hiện chính: ', 'Mô hình trung gian chuỗi (chain mediation): bầu không khí '
      'trường học mà học sinh cảm nhận làm giảm cảm xúc tiêu cực cả trực tiếp lẫn gián tiếp, '
      'qua hai trung gian nối tiếp — cảm giác thuộc về trường (school belonging) và né tránh '
      '– khó chịu xã hội (social avoidance and distress).')
field(doc, 'Truy cập mở: ', 'Có (giấy phép CC BY-NC-ND).')
field(doc, '→ Giá trị cho bài Q1: ', 'Khuôn mẫu SEM trung gian chuỗi ở đúng nhóm tuổi trung '
      'học cơ sở — mẫu phân tích trực tiếp cho đường dẫn nhiều tầng của bài Q1.', color=GREEN)

H(doc, 'B5. Wang và cộng sự (2025) — nghiên cứu dọc, trung gian giấc ngủ/cô đơn', 2)
field(doc, 'Trích dẫn: ', 'Wang, J., Wang, Z., Yang, Y., Wang, T., Lin, H., Zhang, W., '
      'Chen, X., & Fu, C. (2025). Academic burden and emotional problems among adolescents: '
      'A longitudinal mediation analysis. Journal of Adolescence, 97(4), 989–1001.')
field(doc, 'DOI / Nguồn: ', '10.1002/jad.12471 — PMID 39835663 — PMC12128909.')
field(doc, 'Thiết kế & mẫu: ', 'Nghiên cứu thuần tập DỌC 3 đợt (4–5/2022; 9–10/2022; '
      '2–5/2023); 2.965 vị thành niên, tuổi trung bình 15,2; thành phố Taizhou, tỉnh Chiết '
      'Giang, Trung Quốc.')
field(doc, 'Phát hiện chính: ', 'SEM cho thấy gánh nặng học tập dự báo triệu chứng lo âu và '
      'trầm cảm; giấc ngủ và cô đơn là biến trung gian (tỷ trọng trung gian lần lượt 34,1% '
      'và 20,6% cho đường dẫn stress học tập → triệu chứng cảm xúc).')
field(doc, 'Truy cập mở: ', 'Có (giấy phép CC BY).')
field(doc, '→ Giá trị cho bài Q1: ', 'Thiết kế DỌC kết hợp SEM trung gian, gánh nặng học '
      'tập → lo âu — khuôn mẫu cho định hướng nâng cấp bài Q1 lên thiết kế đo lặp; gợi ý bổ '
      'sung trung gian giấc ngủ và cô đơn.', color=GREEN)

H(doc, 'B6. Šutić và cộng sự (2025) — dọc cường độ cao, mô hình SEM động', 2)
field(doc, 'Trích dẫn: ', 'Šutić, L., Yıldız, E., Yavuz Şala, F. C., Duzen, A., Keijsers, '
      'L., & Boele, S. (2025). Parenting and adolescent anxiety within families: A biweekly '
      'longitudinal study. Journal of Child Psychology and Psychiatry, 66(9), 1414–1424.')
field(doc, 'DOI / Nguồn: ', '10.1111/jcpp.14161 — PMID 40104873 — PMC12350824.')
field(doc, 'Thiết kế & mẫu: ', 'Nghiên cứu dọc cường độ cao — đo hai tuần một lần suốt một '
      'năm (26 thời điểm); 256 vị thành niên Hà Lan (tuổi trung bình 14,4) và 176 cha mẹ. '
      'Phân tích bằng Mô hình phương trình cấu trúc động (Dynamic SEM — DSEM).')
field(doc, 'Phát hiện chính: ', 'Khảo sát quan hệ hai chiều giữa thực hành nuôi dạy (hỗ trợ '
      'tự chủ, sự xâm phạm) và triệu chứng lo âu lan tỏa của vị thành niên, ở cấp độ biến '
      'thiên trong nội bộ từng gia đình.')
field(doc, 'Truy cập mở: ', 'Có (giấy phép CC BY).')
field(doc, '→ Giá trị cho bài Q1: ', 'Khuôn mẫu phương pháp dọc cường độ cao kết hợp DSEM — '
      'tách hiệu ứng trong nội bộ cá nhân/gia đình; gợi ý hướng thiết kế tiên tiến. Lưu ý: '
      'mẫu Hà Lan (bối cảnh phương Tây), dùng làm tham chiếu phương pháp.', color=GREEN)

H(doc, 'B7. Yang và cộng sự (2024) — SEM đa trung gian, nuôi dạy hà khắc → lo âu', 2)
field(doc, 'Trích dẫn: ', 'Yang, X., Lin, L., Feng, W., Liu, P., Liang, N., Xue, Z., Ma, '
      'Y., Shen, Y., Yu, W., Lu, J., & Liu, J. (2024). Maternal and paternal harsh '
      'parenting and anxiety symptoms in Chinese adolescents: Examining a multiple '
      'mediation model. Child and Adolescent Psychiatry and Mental Health, 18, Article 134.')
field(doc, 'DOI / Nguồn: ', '10.1186/s13034-024-00826-9 — PMID 39438933 — PMC11515719.')
field(doc, 'Thiết kế & mẫu: ', 'Cắt ngang, SEM kết hợp phân tích đa nhóm; 3.295 vị thành '
      'niên Trung Quốc (54,7% nữ, tuổi trung bình 14,97).')
field(doc, 'Phát hiện chính: ', 'Nuôi dạy hà khắc của mẹ và của cha dự báo triệu chứng lo '
      'âu, qua các trung gian độc lập (gắn kết trường học, nghiện internet, vấn đề giấc '
      'ngủ) và các đường dẫn nối tiếp qua tự hiệu quả (self-efficacy). Hiệu ứng khác nhau '
      'theo hình thức sống chung (cả cha mẹ / chỉ mẹ / chỉ cha).')
field(doc, 'Truy cập mở: ', 'Có (giấy phép CC BY).')
field(doc, '→ Giá trị cho bài Q1: ', 'SEM đa trung gian, nuôi dạy hà khắc → lo âu, mẫu lớn '
      'châu Á — đối chứng trực tiếp cho tầng "gia đình" trong khung đường dẫn; minh hoạ kỹ '
      'thuật phân tích đa nhóm.', color=GREEN)

H(doc, 'B8. Tan và cộng sự (2025) — đường dẫn lan truyền lo âu cha mẹ → con', 2)
field(doc, 'Trích dẫn: ', 'Tan, G., Guo, G., Ouyang, Y., và cộng sự (2025). From parents '
      'to teens: Unpacking the anxiety transmission pathway in parent-adolescent dyads '
      'through parent-adolescent relationships and academic stress. Psychiatric Quarterly. '
      '(Xuất bản trực tuyến trước, 2025).')
field(doc, 'DOI / Nguồn: ', '10.1007/s11126-025-10166-2.')
field(doc, 'Thiết kế & mẫu: ', '2.308 vị thành niên Trung Quốc (tuổi trung bình 15,36; '
      '51,87% nam) và cha mẹ của các em; phân tích bằng SEM.')
field(doc, 'Phát hiện chính: ', 'Lo âu của cha mẹ gián tiếp làm tăng lo âu của vị thành '
      'niên thông qua quan hệ cha mẹ – con bị suy yếu và stress học tập tăng lên; tự khẳng '
      'định bản thân (self-affirmation) mạnh làm dịu tác động tiêu cực của stress học tập '
      'lên lo âu.')
field(doc, 'Truy cập mở: ', 'CHƯA RÕ — Psychiatric Quarterly (Springer) thường là tạp chí '
      'thuê bao; chưa tải được PDF toàn văn. Cần truy cập qua thư viện điện tử của trường.')
field(doc, '→ Giá trị cho bài Q1: ', 'Đường dẫn "lo âu cha mẹ → lo âu con" qua quan hệ và '
      'stress học tập, có yếu tố bảo vệ (tự khẳng định) đóng vai trò điều tiết — khớp khung '
      'đường dẫn nguy cơ – bảo vệ, bối cảnh Á Đông.', color=GREEN)

# ================= GHI CHU =================
H(doc, '4. Ghi chú xác minh và trạng thái PDF', 1)
bullet(doc, 'Cả 8 bài đã được mở trực tiếp trang nguồn (PubMed/PMC/nhà xuất bản) để xác '
            'minh tác giả, năm, tạp chí, số bài và DOI — không trích dẫn theo trí nhớ.', 11)
bullet(doc, 'Đã tải được PDF toàn văn 7/8 bài (A4, A5, A6, B4, B5, B6, B7) — lưu tại thư '
            'mục: bai-bao-khgdvn/bài báo tham khảo/Q1_dot2_19052026/. Mỗi PDF đã kiểm tra '
            'hợp lệ (đầu tệp %PDF, có dấu kết thúc, 9–20 trang).', 11)
bullet(doc, 'Riêng bài B8 (Tan và cs., 2025) chưa tải được — Psychiatric Quarterly là tạp '
            'chí thuê bao. Cần tải qua thư viện điện tử; liên kết tra cứu: '
            'https://doi.org/10.1007/s11126-025-10166-2', 11)
bullet(doc, 'Tám bài này MỚI là ứng viên tham khảo cho bài Q1 — chưa nhập vào thư viện '
            'canonical của dự án (chưa có mã VN/QT chính thức, chưa có bản dịch/tóm tắt). '
            'Nếu quyết định dùng, cần chạy bước nhập canonical như các tài liệu khác.', 11)
bullet(doc, 'Khoảng trống ghi nhận khi tìm kiếm: nghiên cứu Việt Nam dùng thiết kế DỌC '
            'hoặc mô hình SEM/đường dẫn cho LO ÂU ở học sinh trung học còn rất ít — phần '
            'lớn là cắt ngang. Đây chính là khoảng trống mà bài Q1 có thể khai thác để lập '
            'luận về tính mới.', 11)

H(doc, '5. Tham khảo & truy vết', 1)
P(doc, 'Liên kết kiểm chứng (DOI — mở trực tiếp tới bài gốc):', bold=True, size=11)
for label, url in [
    ('A4 — doi.org/10.52965/001c.116967', 'https://doi.org/10.52965/001c.116967'),
    ('A5 — doi.org/10.1017/gmh.2025.10096', 'https://doi.org/10.1017/gmh.2025.10096'),
    ('A6 — doi.org/10.1007/s12144-022-03661-3', 'https://doi.org/10.1007/s12144-022-03661-3'),
    ('B4 — doi.org/10.1186/s40359-025-02364-1', 'https://doi.org/10.1186/s40359-025-02364-1'),
    ('B5 — doi.org/10.1002/jad.12471', 'https://doi.org/10.1002/jad.12471'),
    ('B6 — doi.org/10.1111/jcpp.14161', 'https://doi.org/10.1111/jcpp.14161'),
    ('B7 — doi.org/10.1186/s13034-024-00826-9', 'https://doi.org/10.1186/s13034-024-00826-9'),
    ('B8 — doi.org/10.1007/s11126-025-10166-2', 'https://doi.org/10.1007/s11126-025-10166-2'),
]:
    p = doc.add_paragraph(style='List Bullet'); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    add_hyperlink(p, url, label, size=11)

P(doc, 'Truy vết nội bộ:', bold=True, size=11)
bullet(doc, 'PDF toàn văn: bai-bao-khgdvn/bài báo tham khảo/Q1_dot2_19052026/ (7 tệp).', 11)
bullet(doc, 'Script tìm – tải: 06_Scripts/tai_pdf_Q1_dot2_19052026.py và '
            'tai_pdf_Q1_dot2_retry_19052026.py (tải qua Europe PMC).', 11)
bullet(doc, 'Doc đợt 1 (6 bài): bai-bao-khgdvn/Bài tham khảo đã xác minh - chuẩn bị bài '
            'Q1.docx — tài liệu này nối tiếp, dùng mã A4–A6, B4–B8.', 11)
P(doc, 'Lưu ý: A6 (Ho và cs., 2022) đo kết cục trầm cảm; B6 (Šutić và cs., 2025) dùng mẫu '
       'Hà Lan. Khi đưa vào bài Q1 cần nêu rõ hai điểm này để tránh suy luận quá mức.',
  italic=True, size=11)

doc.save(OUT)
print('DA LUU:', OUT)
print('So bang:', len(doc.tables), '| So doan:', len(doc.paragraphs))
