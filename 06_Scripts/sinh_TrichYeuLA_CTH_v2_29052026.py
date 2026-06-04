# -*- coding: utf-8 -*-
"""Sinh Trich yeu LA cho NCS Cong Thi Hang - v2 sau khi fact-check + format audit.

FIXES tu v1:
- Fact: MPAS -> SAS-SV (Kwon 2013); MPVS -> OBVQ (Olweus 1996); Rosenberg -> RSES
- Fact: '~86%' -> '~85-89%' (khop LA P1416)
- Fact: '9310401' -> '9.31.04.01' (khop LA chinh P48)
- Format: label bold + value KHONG bold (khop template runs)
- Format: line spacing 1.5 (khop template, KHONG phai 1.15)
- Format: first line indent body 1.0cm (khop template)
- Format: title 'TRICH YEU LA' space_before 11.7pt
- Format: 'New conclusions' space_before/after 3pt

29/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', 'TrichYeuLA_CongThiHang_v2_29052026.docx')

d = Document()

# PAGE SETUP - khop template
for sec in d.sections:
    sec.page_width = Cm(21.0)
    sec.page_height = Cm(29.7)
    sec.top_margin = Cm(2.54)
    sec.bottom_margin = Cm(2.54)
    sec.left_margin = Cm(2.54)
    sec.right_margin = Cm(2.54)

# Style Normal
s = d.styles['Normal']
s.font.name = 'Times New Roman'
s.font.size = Pt(13)
s.paragraph_format.line_spacing = 1.5
s.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


# ============================================================
# HELPERS
# ============================================================
def title_center(text, space_before_pt=11.7):
    """TRICH YEU LA / DISSERTATION ABSTRACT title."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if space_before_pt:
        p.paragraph_format.space_before = Pt(space_before_pt)
    pf = p.paragraph_format
    pf.line_spacing = 1.0
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = True
    return p


def field(label, value):
    """Header field: LABEL bold + VALUE not bold (khop template)."""
    p = d.add_paragraph()
    p.alignment = None  # default LEFT
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE

    # Label run - BOLD
    r1 = p.add_run(label)
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    r1.bold = True

    # Value run - NOT bold
    r2 = p.add_run(value)
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(13)
    r2.bold = False
    return p


def heading_center(text, bold=True, italic=False, space_before=None, space_after=None):
    """Heading 'Nhung ket luan moi' centered."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if space_before:
        pf.space_before = Pt(space_before)
    if space_after:
        pf.space_after = Pt(space_after)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = bold
    r.italic = italic
    return p


def heading_justify(text, bold=True, italic=False):
    """Heading 'Dong gop ly luan' justified with first line indent."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = bold
    r.italic = italic
    return p


def body(text):
    """Body paragraph justified, first line indent 1cm, NOT bold."""
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    pf = p.paragraph_format
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    pf.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.bold = False
    return p


def signature_table(left_lines, right_lines):
    t = d.add_table(rows=1, cols=2)
    t.autofit = True
    cell_left = t.rows[0].cells[0]
    cell_right = t.rows[0].cells[1]
    cell_left.text = ''
    cell_right.text = ''

    for i, line in enumerate(left_lines):
        p = cell_left.paragraphs[0] if i == 0 else cell_left.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        r = p.add_run(line)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.bold = bool(i == 0 or (line.strip() and i == len(left_lines) - 1))

    for i, line in enumerate(right_lines):
        p = cell_right.paragraphs[0] if i == 0 else cell_right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf = p.paragraph_format
        pf.line_spacing = 1.5
        r = p.add_run(line)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(13)
        r.bold = bool(i == 0 or (line.strip() and i == len(right_lines) - 1))
    return t


# ============================================================
# VIETNAMESE SECTION
# ============================================================

title_center('TRÍCH YẾU LUẬN ÁN', space_before_pt=11.7)
d.add_paragraph()

# Fields - LABEL bold + VALUE not bold (khop template)
field('Tên đề tài: ', 'Rối loạn lo âu ở học sinh trung học cơ sở')
field('Chuyên ngành: ', 'Tâm lý học chuyên ngành')
field('Mã số: ', '9.31.04.01')
field('Nghiên cứu sinh: ', '\t\tCông Thị Hằng')
field('Cán bộ hướng dẫn: ', '\t\tTS. Đào Minh Đức')
field('Cơ sở đào tạo: ', 'Trường Đại học Sư phạm Hà Nội')

# Heading centered: Nhung ket luan moi
heading_center('Những kết luận mới của luận án')

# Dong gop ly luan (heading justified bold)
heading_justify('Đóng góp về mặt lý luận', bold=True)

body('Luận án đã hệ thống hóa và làm rõ cơ sở lý luận về rối loạn lo âu ở học sinh '
     'trung học cơ sở theo hướng tiếp cận đa chiều, bao gồm ba dạng rối loạn lo âu '
     'phổ biến nhất ở trẻ vị thành niên theo phân loại DSM-5: rối loạn lo âu lan tỏa, '
     'rối loạn lo âu xã hội và rối loạn lo âu chia ly. Đồng thời, nghiên cứu xác lập '
     'rối loạn lo âu như một hệ thống đa thành phần, bao gồm các khía cạnh nhận thức, '
     'cảm xúc, sinh lý và hành vi, qua đó làm rõ đặc điểm tâm lý đặc thù của lứa tuổi '
     'học sinh trung học cơ sở trong giai đoạn phát triển dậy thì — giai đoạn có nhiều '
     'biến đổi về nhận thức bản thân và quan hệ xã hội.')

body('Bên cạnh đó, luận án xác lập mô hình mối quan hệ giữa rối loạn lo âu với hệ '
     'thống các yếu tố ảnh hưởng, bao gồm ba nhóm yếu tố nguy cơ (áp lực học tập, '
     'nghiện điện thoại, bắt nạt học đường), bốn nhóm yếu tố bảo vệ (lòng tự trọng, '
     'gắn bó với trường học, hỗ trợ của cha mẹ, hỗ trợ của bạn bè) và biện pháp đối '
     'phó của học sinh. Đóng góp này góp phần lấp đầy khoảng trống nghiên cứu về '
     'phương pháp luận trong lĩnh vực sức khỏe tâm thần học đường ở Việt Nam, đặc '
     'biệt ở quy trình kiểm định chất lượng công cụ đo lường và chuẩn hóa thang đo '
     'phù hợp với ngôn ngữ, văn hóa Việt Nam.')

# Dong gop thuc tien (heading justified bold)
heading_justify('Đóng góp về mặt thực tiễn', bold=True)

body('Về phương pháp đo lường, luận án đã hiệu chỉnh và kiểm định độ tin cậy, độ giá '
     'trị của tám thang đo tâm lý — RCADS (Revised Children\'s Anxiety and Depression '
     'Scale; Chorpita, 2000), ESSA (Educational Stress Scale for Adolescents; Sun và '
     'cộng sự, 2011), SAS-SV (Smartphone Addiction Scale - Short Version; Kwon và cộng '
     'sự, 2013), OBVQ (Olweus Bully/Victim Questionnaire; Olweus, 1996), PSSM '
     '(Psychological Sense of School Membership; Goodenow, 1993), MSPSS '
     '(Multidimensional Scale of Perceived Social Support; Zimet và cộng sự, 1988), '
     'RSES (Rosenberg Self-Esteem Scale; Rosenberg, 1965) và Brief COPE '
     '(Carver, 1997) — thông qua các chỉ số Cronbach\'s alpha, McDonald\'s omega và '
     'phân tích nhân tố khẳng định (CFA). Các thang đo đều đạt ngưỡng chấp nhận, tạo '
     'ra bộ công cụ đo lường có độ ổn định và độ tin cậy cao, phù hợp với bối cảnh '
     'học sinh trung học cơ sở tại Việt Nam.')

body('Kết quả nghiên cứu thực trạng trên mẫu 1.352 học sinh trung học cơ sở tại Hà Nội '
     'đã kiểm chứng thành công ba giả thuyết khoa học: (1) học sinh nữ có mức độ rối '
     'loạn lo âu lan tỏa và rối loạn lo âu xã hội cao hơn học sinh nam (p < 0,001), '
     'trong khi rối loạn lo âu chia ly không có khác biệt theo giới tính (F = 0,246; '
     'p = 0,620); (2) lòng tự trọng là yếu tố bảo vệ có ảnh hưởng mạnh đến rối loạn '
     'lo âu, với cường độ tác động tương đương khoảng 85–89% so với áp lực học tập '
     'trong mô hình phương trình cấu trúc (SEM); (3) ba nhóm yếu tố nguy cơ có thứ tự '
     'cường độ từ cao đến thấp là: áp lực học tập (ĐTB = 51,13) > nghiện điện thoại '
     '(ĐTB = 28,38) > bắt nạt thể chất (ĐTB = 13,52). Bên cạnh đó, kết quả mô hình '
     'SEM cho thấy các chiến lược đối phó hiện tại của học sinh trung học cơ sở chưa '
     'thực sự hiệu quả, thể hiện qua xu hướng gia tăng triệu chứng rối loạn lo âu khi '
     'mức độ sử dụng các biện pháp đối phó tăng — gợi ý sự cần thiết định hướng học '
     'sinh chuyển từ phản ứng né tránh sang các chiến lược chủ động như giải quyết '
     'vấn đề, điều chỉnh nhận thức và tìm kiếm hỗ trợ xã hội.')

body('Trên cơ sở các phát hiện thực nghiệm, luận án đề xuất Khung chương trình tập '
     'huấn phòng ngừa rối loạn lo âu cho học sinh trung học cơ sở gồm 8 nội dung cơ '
     'bản, đáp ứng yêu cầu cấp thiết hiện nay về chăm sóc sức khỏe tâm thần học đường '
     'ở Việt Nam. Khung chương trình này cung cấp cơ sở khoa học cho các kiến nghị '
     'can thiệp đối với học sinh, giáo viên và nhà trường, gia đình, cùng các lực '
     'lượng hỗ trợ học đường (chuyên viên tâm lý, phòng tư vấn học đường, đoàn thanh '
     'niên, các câu lạc bộ); đồng thời là tài liệu tham khảo có giá trị cho giáo '
     'viên, chuyên viên tâm lý học đường và nhà quản lý giáo dục trong công tác phòng '
     'ngừa và hỗ trợ học sinh có biểu hiện rối loạn lo âu.')

d.add_paragraph()

signature_table(
    left_lines=['Người hướng dẫn', '', '', '', 'TS. Đào Minh Đức'],
    right_lines=['Nghiên cứu sinh', '', '', '', 'Công Thị Hằng'],
)

d.add_page_break()


# ============================================================
# ENGLISH SECTION
# ============================================================

title_center('DISSERTATION ABSTRACT', space_before_pt=11.7)
d.add_paragraph()

field('Title: ', 'Anxiety Disorders among Lower Secondary School Students')
field('Speciality: ', 'Psychology')
field('Classification: ', '9.31.04.01')
field('Name of PhD Student: ', '\t\tCong Thi Hang')
field('Advisor: ', '\t\tDr. Dao Minh Duc')
field('Institutional: ', 'Hanoi National University of Education')

# Heading centered + space before/after 3pt
heading_center('New conclusions', space_before=3, space_after=3)

# Theoretical contribution - bold + ITALIC + justify (khop template)
heading_justify('Theoretical contribution', bold=True, italic=True)

body('The dissertation has systematized and clarified the theoretical foundations of '
     'anxiety disorders among lower secondary school students through a multidimensional '
     'approach, encompassing the three most prevalent anxiety disorder types in '
     'adolescents according to the DSM-5 classification: generalized anxiety disorder, '
     'social anxiety disorder, and separation anxiety disorder. Furthermore, the study '
     'conceptualizes anxiety disorders as a multi-component system comprising cognitive, '
     'emotional, physiological, and behavioral dimensions, thereby illuminating the '
     'distinctive psychological characteristics of lower secondary school students '
     'during the pubertal developmental stage — a period marked by substantial '
     'transformations in self-perception and social relationships.')

body('In addition, the dissertation establishes a model of the relationships between '
     'anxiety disorders and a system of influencing factors, including three risk '
     'factor groups (academic pressure, smartphone addiction, school bullying), four '
     'protective factor groups (self-esteem, school engagement, parental support, peer '
     'support), and students\' coping strategies. This contribution helps fill the '
     'methodological gap in the field of school-based mental health research in '
     'Vietnam, particularly regarding the validation procedures for measurement '
     'instruments and the standardization of psychological scales to align with the '
     'Vietnamese language and cultural context.')

# Practical contributions - bold + ITALIC + justify
heading_justify('Practical contributions', bold=True, italic=True)

body('Methodologically, the dissertation has adapted and validated the reliability and '
     'validity of eight psychological scales — RCADS (Revised Children\'s Anxiety and '
     'Depression Scale; Chorpita, 2000), ESSA (Educational Stress Scale for '
     'Adolescents; Sun et al., 2011), SAS-SV (Smartphone Addiction Scale - Short '
     'Version; Kwon et al., 2013), OBVQ (Olweus Bully/Victim Questionnaire; Olweus, '
     '1996), PSSM (Psychological Sense of School Membership; Goodenow, 1993), MSPSS '
     '(Multidimensional Scale of Perceived Social Support; Zimet et al., 1988), RSES '
     '(Rosenberg Self-Esteem Scale; Rosenberg, 1965), and Brief COPE (Carver, 1997) — '
     'through Cronbach\'s alpha, McDonald\'s omega, and confirmatory factor analysis '
     '(CFA). All scales attained acceptable thresholds, establishing a measurement '
     'toolkit with high stability and reliability suitable for the lower secondary '
     'school context in Vietnam.')

body('The empirical findings from a sample of 1,352 lower secondary school students in '
     'Hanoi successfully confirmed three scientific hypotheses: (1) female students '
     'exhibit higher levels of generalized anxiety disorder and social anxiety '
     'disorder than male students (p < 0.001), whereas separation anxiety shows no '
     'significant gender differences (F = 0.246; p = 0.620); (2) self-esteem serves '
     'as a protective factor with a strong influence on anxiety disorders, with a '
     'magnitude approximately 85–89% comparable to academic pressure within the '
     'structural equation model (SEM); (3) the three risk factor groups, in descending '
     'order of magnitude, are: academic pressure (M = 51.13) > smartphone addiction '
     '(M = 28.38) > physical bullying (M = 13.52). Additionally, the SEM results '
     'reveal that current coping strategies of lower secondary school students are not '
     'yet sufficiently effective, evidenced by the upward trend in anxiety symptoms '
     'accompanying increased use of coping behaviors — suggesting the need to guide '
     'students from avoidance reactions toward proactive strategies such as '
     'problem-solving, cognitive reappraisal, and social support seeking.')

body('Building upon these empirical findings, the dissertation proposes an '
     'Eight-Component Framework for the Prevention of Anxiety Disorders among lower '
     'secondary school students, addressing the urgent contemporary need for '
     'school-based mental health care in Vietnam. This framework provides a scientific '
     'basis for intervention recommendations targeting students, teachers and schools, '
     'families, and school support personnel (school psychologists, school counseling '
     'offices, youth unions, student clubs); it also serves as a valuable reference '
     'document for teachers, school psychologists, and education administrators in the '
     'prevention of and support for students manifesting anxiety symptoms.')

d.add_paragraph()

signature_table(
    left_lines=['Advisor', '', '', '', 'Dr. Dao Minh Duc'],
    right_lines=['PhD Candidate', '', '', '', 'Cong Thi Hang'],
)


# Clean metadata
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
cp.identifier = ''
cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'File size: {os.path.getsize(OUT)} bytes')
