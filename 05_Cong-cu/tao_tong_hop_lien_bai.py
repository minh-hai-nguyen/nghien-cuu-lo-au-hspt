# -*- coding: utf-8 -*-
"""Tổng hợp liên bài báo — Phân tích sâu cross-study synthesis"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(6)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def h(text, level=1):
    hd = doc.add_heading(text, level=level)
    for r in hd.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def p(text, bold=False, italic=False, size=12):
    para = doc.add_paragraph()
    r = para.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return para

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w * 567)))
    we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, hdr in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = hdr
        for pp in c.paragraphs:
            pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in pp.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(9)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for pp in c.paragraphs:
                for r in pp.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(9)
    return t

# ===== BÌA =====
title = doc.add_heading('', level=0)
r = title.add_run('TỔNG HỢP LIÊN BÀI BÁO\nLO ÂU Ở HỌC SINH THCS VÀ THPT\nPhân tích sâu từ 10 nghiên cứu (2022\u20132025)')
r.font.name = 'Times New Roman'
r.font.size = Pt(18)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
p('Việt Nam \u2014 Đông Nam Á \u2014 Thế giới | Tháng 3/2026', italic=True, size=11)
doc.add_page_break()

# ===== 1. TỶ LỆ THEO QUỐC GIA =====
h('1. Tỷ lệ lo âu và trầm cảm theo quốc gia, công cụ và năm')
p('Bảng dưới đây tổng hợp tỷ lệ lo âu/trầm cảm từ 8 nghiên cứu tại 5 quốc gia/khu vực. Điểm quan trọng nhất là sự khác biệt rất lớn giữa các nghiên cứu — phản ánh sự khác nhau về công cụ đo, ngưỡng cắt, mẫu và thời điểm, chứ không nhất thiết phản ánh sự khác nhau về mức độ lo âu thực sự.')
tbl(
    ['Quốc gia/Vùng', 'Địa bàn', 'Tỷ lệ', 'Công cụ', 'n', 'Đối tượng', 'Năm KS', 'Nguồn'],
    [
        ['Việt Nam', 'Hà Nội', 'Lo âu: 40,6%', 'GAD-7 (sàng lọc)', '3.910', 'HS THPT 14-17t', '10-11/2021', 'Hoa 2024'],
        ['Việt Nam', '38 tỉnh (quốc gia)', 'Lo âu: 18,6% (vấn đề)\n2,3% (rối loạn DSM-5)', 'DISC-5 (chẩn đoán)', '5.996', 'VTN 10-17t', '2021-22', 'V-NAMHS'],
        ['Việt Nam', 'Lạng Sơn (DTTS)', 'Lo âu: 54,4%\nTrầm cảm: 59,0%', 'DASS-21 (sàng lọc)', '845', 'HS nội trú DTTS', '2024', 'Ngô Anh Vinh'],
        ['Việt Nam', 'Huế (SSF)', 'Lo âu: M=22,96\nβ cảm xúc=\u22120,40', 'Thang VN + CES-D', '273+273', 'VTN 12-18t SSF', '2024', 'Pham 2024'],
        ['Trung Quốc', 'Suzhou (Đông)', 'Trầm cảm: 14,5%+5,8%', 'PHQ-9 (sàng lọc)', '9.831', 'HS THCS/THPT', '2019-23', 'Zhu 2025'],
        ['Philippines', 'Quốc gia', 'Trầm cảm: 9,6%\u219220,9%', 'CES-D-11 (sàng lọc)', '30.127', 'TN 15-24t', '2013\u219221', 'Puyat 2025'],
        ['ASEAN', '10 nước', 'RLTT chung: 11,9%\nVN: 10,1%', 'GBD/DisMod-MR', 'Dân số', 'Tất cả tuổi', '1990-2021', 'GBD 2025'],
        ['Nam Á', '8 nước', '1,5% \u2013 81,6%', 'Đa dạng', 'Tổng quan', 'VTN 10-19t', '\u22642024', 'Mudunna 2025'],
    ],
    widths=[1.8, 1.5, 2.2, 1.5, 1.0, 1.3, 1.0, 1.2]
)

doc.add_paragraph()
p('Phân tích:', bold=True)
p('a) Tỷ lệ dao động từ 2,3% đến 81,6% — khoảng cách gấp 35 lần. Nguyên nhân chính không phải sự khác biệt về mức độ lo âu giữa các nước, mà là sự khác nhau về công cụ đo. Thang sàng lọc (DASS-21, GAD-7, CES-D) đo TRIỆU CHỨNG, cho tỷ lệ cao; công cụ chẩn đoán (DISC-5/DSM-5) đo RỐI LOẠN ĐẦY ĐỦ, cho tỷ lệ thấp hơn 10-37 lần (COVID-19 Mental Disorders Collaborators, 2021; Ferrari và cộng sự, 2013).')
p('b) Ngay trong cùng một quốc gia (Việt Nam), tỷ lệ dao động từ 2,3% (V-NAMHS, DISC-5) đến 54,4% (Ngô Anh Vinh, DASS-21). Đây là bằng chứng mạnh nhất cho thấy KHÔNG THỂ so sánh trực tiếp các nghiên cứu dùng công cụ khác nhau.')
p('c) Nhóm dễ bị tổn thương có tỷ lệ cao bất thường: DTTS Lạng Sơn 54,4% (Ngô Anh Vinh, 2024), VTN SSF Huế lo âu M=22,96 (Pham, 2024), LGBTQ+ Philippines 32,3% (Puyat, 2025).')

doc.add_page_break()

# ===== 2. SÀNG LỌC VS CHẨN ĐOÁN =====
h('2. Khoảng cách sàng lọc \u2014 chẩn đoán: Bằng chứng liên bài')
p('Đây là phát hiện quan trọng nhất khi tổng hợp các bài. V-NAMHS (2022) là nghiên cứu duy nhất tại Việt Nam dùng công cụ chẩn đoán chuẩn hóa (DISC-5/DSM-5), cho tỷ lệ rất khác với sàng lọc.')
tbl(
    ['Phương pháp', 'Tỷ lệ lo âu VN', 'Công cụ', 'n', 'Nguồn'],
    [
        ['Sàng lọc (screening)', '40,6%', 'GAD-7', '3.910', 'Hoa 2024 (Hà Nội)'],
        ['Sàng lọc (screening)', '54,4%', 'DASS-21', '845', 'Ngô Anh Vinh 2024 (DTTS)'],
        ['Chẩn đoán (diagnostic)', '2,3%', 'DISC-5/DSM-5', '5.996', 'V-NAMHS 2022 (quốc gia)'],
        ['Chênh lệch', 'Gấp 17-24 lần', '', '', 'So sánh liên bài'],
    ],
    widths=[2.5, 2.0, 2.0, 1.5, 3.5]
)
doc.add_paragraph()
p('Hệ quả cho nghiên cứu:', bold=True)
p('\u2022 Khi trích dẫn tỷ lệ lo âu, BẮT BUỘC phải ghi rõ công cụ sử dụng và đây là sàng lọc hay chẩn đoán.')
p('\u2022 Tỷ lệ sàng lọc 40-86% KHÔNG CÓ NGHĨA 40-86% HS có rối loạn lo âu — chỉ có triệu chứng ở mức sàng lọc.')
p('\u2022 Cần nghiên cứu so sánh cả 2 phương pháp trên CÙNG MẪU để xác định ngưỡng cắt phù hợp cho VN (gap).')

doc.add_page_break()

# ===== 3. GIỚI TÍNH =====
h('3. Yếu tố giới tính — Nhất quán xuyên suốt các NC')
p('Đây là yếu tố nguy cơ NHẤT QUÁN NHẤT — nữ giới có tỷ lệ lo âu/trầm cảm cao hơn nam giới trong TẤT CẢ nghiên cứu có phân tích theo giới (trừ V-NAMHS không tìm thấy khác biệt có ý nghĩa — có thể do DISC-5 không nhạy với biểu hiện theo giới ở VN).')
tbl(
    ['Nghiên cứu', 'Quốc gia', 'Nữ', 'Nam', 'Chỉ số', 'Ý nghĩa'],
    [
        ['Hoa 2024', 'VN (Hà Nội)', 'GAD-7 M=1,74', 'M=1,50', 'p < 0,01', 'Nữ cao hơn 16%'],
        ['Zhu 2025', 'TQ (Suzhou)', 'AOR=1 (ref)', 'AOR=0,803', 'p < 0,05', 'Nam bảo vệ 20%'],
        ['Puyat 2025', 'Philippines', '24,3%', '17,0%', 'aPR=1,41', 'Nữ cao hơn 41%'],
        ['Anderson 2025', 'Thế giới', 'Cao hơn', 'Thấp hơn', 'Nhất quán', 'Tổng quan 61 bài'],
        ['V-NAMHS', 'VN (quốc gia)', '22,6%', '20,8%', 'Không ý nghĩa', 'Ngoại lệ — cần giải thích'],
    ],
    widths=[2.0, 2.0, 2.0, 2.0, 2.0, 2.5]
)
doc.add_paragraph()
p('Phân tích liên bài:', bold=True)
p('Sự chênh lệch giới dao động từ 16% (VN, Hoa 2024) đến 41% (Philippines, Puyat 2025). Điều này gợi ý rằng yếu tố văn hóa xã hội — không chỉ sinh học — đóng vai trò quan trọng. Philippines có tỷ lệ cao nhất có thể liên quan đến bạo lực trên cơ sở giới và áp lực xã hội đặc thù (McLean và cộng sự, 2011; Daly, 2022).')
p('Ngoại lệ V-NAMHS: Không tìm thấy khác biệt giới tính có ý nghĩa (nam 20,8% vs nữ 22,6%). Có thể do DISC-5 chưa được chuẩn hóa cho văn hóa VN — nam giới VN có thể biểu hiện triệu chứng khác (somatization thay vì cảm xúc) mà DISC-5 không đo được (Canino & Alegria, 2008).')

doc.add_page_break()

# ===== 4. GIẤC NGỦ =====
h('4. Giấc ngủ — Yếu tố nguy cơ mạnh nhất cho trầm cảm')
p('Phát hiện nổi bật nhất từ Zhu 2025 (Trung Quốc, n=9.831): ngủ dưới 5 giờ tăng nguy cơ trầm cảm chắc chắn lên 13,7 lần (AOR=13,710). Đây là yếu tố nguy cơ có AOR cao nhất trong TẤT CẢ 10 bài nghiên cứu.')
tbl(
    ['Giờ ngủ', 'AOR (trầm cảm có thể)', 'AOR (trầm cảm chắc chắn)', 'Nguồn'],
    [
        ['< 5 giờ', '3,901', '13,710', 'Zhu 2025'],
        ['5-6 giờ', '3,759', '7,289', 'Zhu 2025'],
        ['6-7 giờ', '1,961', '2,678', 'Zhu 2025'],
        ['7-8 giờ', '1,372', '1,372', 'Zhu 2025'],
        ['\u2265 8 giờ', '1 (ref)', '1 (ref)', 'Zhu 2025'],
    ],
    widths=[3.0, 3.0, 3.0, 2.5]
)
doc.add_paragraph()
p('So sánh liên bài:', bold=True)
p('\u2022 Hoa 2024 (VN): không đo giấc ngủ trực tiếp nhưng phỏng vấn cho thấy "lo lắng ảnh hưởng giấc ngủ" là chủ đề nổi bật.')
p('\u2022 Hoàng Trung Học 2025 (VN): "chất lượng giấc ngủ kém" là một trong 3 yếu tố nguy cơ chính.')
p('\u2022 Anderson 2025 (thế giới): tổng quan ghi nhận mạng xã hội \u2192 thiếu ngủ \u2192 lo âu là chuỗi nhân quả phổ biến (Twenge và cộng sự, 2018; Hale và cộng sự, 2015).')
p('Gap: Chưa có NC nào tại Việt Nam đo lường mối quan hệ giấc ngủ \u2014 lo âu với AOR cụ thể như Zhu 2025 làm tại TQ. Đây là hướng NC quan trọng.', italic=True)

doc.add_page_break()

# ===== 5. CẤU TRÚC GIA ĐÌNH =====
h('5. Cấu trúc gia đình và chăm sóc cảm xúc — Yếu tố bảo vệ xuyên văn hóa')
p('Bằng chứng từ nhiều bài cho thấy gia đình là yếu tố bảo vệ (hoặc nguy cơ) mạnh nhất, xuyên suốt các nền văn hóa:')
tbl(
    ['Yếu tố gia đình', 'Tác động', 'Chỉ số', 'Nguồn'],
    [
        ['Gia đình đơn thân', 'Nguy cơ', 'AOR = 1,434', 'Zhu 2025 (TQ)'],
        ['Gia đình tái hôn', 'Nguy cơ', 'AOR = 1,837', 'Zhu 2025 (TQ)'],
        ['Cả 2 cha mẹ vắng mặt', 'Nguy cơ', 'AOR = 1,710', 'Zhu 2025 (TQ)'],
        ['Quan hệ cha mẹ-con kém', 'Nguy cơ', 'Yếu tố chính', 'Hoàng TH 2025 (VN)'],
        ['Chăm sóc cảm xúc tốt', 'Bảo vệ (mạnh)', '\u03b2 = \u22120,40 cho lo âu', 'Pham 2024 (VN)'],
        ['Chăm sóc thể chất', 'Không ý nghĩa', '\u03b2 không ý nghĩa', 'Pham 2024 (VN)'],
        ['Có người tâm sự', 'Bảo vệ (rất mạnh)', 'OR = 0,22 cho trầm cảm', 'Dong 2025 (TQ)'],
        ['Kỳ vọng gia đình', 'Nguy cơ', 'Chủ đề định tính', 'Hoa 2024 (VN)'],
        ['ACEs (trải nghiệm bất lợi)', 'Nguy cơ', 'Hệ số 0,28 với lo âu', 'Ngô Anh Vinh (VN)'],
    ],
    widths=[3.0, 1.5, 3.0, 3.0]
)
doc.add_paragraph()
p('Phát hiện liên bài đặc biệt quan trọng:', bold=True)
p('Pham 2024 (VN, Huế) cho thấy chăm sóc CẢM XÚC quan trọng hơn chăm sóc THỂ CHẤT (\u03b2 = \u22120,40 vs không ý nghĩa). Phát hiện này tương đồng với Dong 2025 (TQ) khi "có người tâm sự" giảm 78% nguy cơ trầm cảm (OR = 0,22). Cả hai gợi ý rằng can thiệp nên tập trung vào CHẤT LƯỢNG MỐI QUAN HỆ hơn là điều kiện vật chất.')
p('Gap: Chưa có RCT nào đào tạo cha mẹ/người chăm sóc VN về kỹ năng chăm sóc cảm xúc \u2192 đo SKTT con (Werner-Seidler và cộng sự, 2021).', italic=True)

doc.add_page_break()

# ===== 6. XU HƯỚNG THỜI GIAN =====
h('6. Xu hướng thời gian — Lo âu/trầm cảm đang tăng hay giảm?')
tbl(
    ['Khu vực', 'Giai đoạn', 'Xu hướng', 'Bằng chứng', 'Nguồn'],
    [
        ['ASEAN 10 nước', '1990\u21922021', 'TĂNG 70%', '47,3M \u2192 80,4M ca RLTT', 'GBD 2025'],
        ['Philippines', '2013\u21922021', 'TĂNG GẤP ĐÔI', 'Trầm cảm 9,6% \u2192 20,9%', 'Puyat 2025'],
        ['Việt Nam', 'Trong \u2192 Sau COVID', 'GIẢM nhưng vẫn CAO', 'Lo âu 41,5% \u2192 25,4%', 'Hoàng TH 2025'],
        ['Trung Quốc', '2019\u21922023', 'ỔN ĐỊNH ~20%', 'Trầm cảm 14,5%+5,8%', 'Zhu 2025'],
        ['Thế giới', '2020\u21922025', 'TĂNG (Gen Z cao nhất)', '31,9% VTN có lo âu', 'Anderson 2025'],
    ],
    widths=[2.0, 2.0, 2.0, 3.0, 2.0]
)
doc.add_paragraph()
p('Phân tích xu hướng:', bold=True)
p('a) Xu hướng dài hạn (30 năm): Rõ ràng TĂNG — GBD ASEAN cho thấy +70% số ca từ 1990 đến 2021. Xu hướng này nhất quán với dữ liệu toàn cầu (Polanczyk và cộng sự, 2015).')
p('b) Tác động COVID-19: Tăng mạnh TRONG dịch, giảm SAU dịch nhưng không về mức trước dịch. Tại VN, lo âu giảm 16,1 điểm % (41,5% \u2192 25,4%) nhưng 25,4% vẫn cao (Hoàng TH 2025). Philippines tăng mạnh nhất: gấp đôi trong 8 năm (Puyat 2025).')
p('c) Thế hệ Z: Anderson 2025 xác nhận Gen Z (sinh 1997-2012) có tỷ lệ lo âu cao nhất trong 3 thế hệ — liên quan mạng xã hội, áp lực học tập, bất ổn chính trị.')
p('Gap: Chưa có NC dọc nào tại VN theo dõi CÙNG NHÓM HS qua nhiều năm (trước-trong-sau COVID).', italic=True)

doc.add_page_break()

# ===== 7. CAN THIỆP =====
h('7. Can thiệp tại trường học — Bằng chứng còn rất yếu')
p('Zhameden 2025 là tổng quan hệ thống duy nhất về can thiệp trường học tại LMIC. Kết quả đáng lo ngại:')
tbl(
    ['Tiêu chí', 'Trầm cảm', 'Lo âu', 'Ghi chú'],
    [
        ['Số NC có hiệu quả', '3/4 (75%)', '1/4 (25%)', 'Can thiệp kém hiệu quả cho lo âu'],
        ['Phương pháp chính', 'CBT (4/6)', 'CBT', 'Thiếu đa dạng'],
        ['Nước nghiên cứu', 'Brazil, TQ, Ấn Độ, Kenya, Lebanon, Malaysia', '', 'KHÔNG CÓ Việt Nam'],
        ['Chất lượng bằng chứng', 'RẤT THẤP (GRADE)', '', 'Tất cả high risk of bias'],
        ['Theo dõi dài nhất', '3 tháng', '', 'Quá ngắn'],
    ],
    widths=[3.0, 3.0, 2.5, 3.0]
)
doc.add_paragraph()
p('So sánh với bằng chứng khác:', bold=True)
p('\u2022 Anderson 2025 khuyến nghị can thiệp đa tầng (lâm sàng + giáo dục + cộng đồng) — không chỉ CBT.')
p('\u2022 Pham 2024 gợi ý can thiệp nâng cao chăm sóc cảm xúc cho nhân viên SSF có thể hiệu quả hơn CBT đơn thuần.')
p('\u2022 Zhu 2025 cho thấy can thiệp giấc ngủ + hoạt động ngoài trời có thể là phương pháp "tự nhiên" hiệu quả.')
p('Gap lớn nhất: CHƯA CÓ RCT NÀO tại Việt Nam về can thiệp SKTT tại trường học. Đây là khoảng trống bằng chứng cấp thiết nhất.', italic=True)

doc.add_page_break()

# ===== 8. DỊCH VỤ SKTT =====
h('8. Khoảng trống dịch vụ SKTT — "Biết bệnh nhưng không đến bác sĩ"')
tbl(
    ['Chỉ số', 'Giá trị', 'Quốc gia', 'Nguồn'],
    [
        ['VTN có vấn đề SKTT tiếp cận dịch vụ', 'Chỉ 8,4%', 'Việt Nam', 'V-NAMHS 2022'],
        ['Phụ huynh nhận ra con cần giúp đỡ', 'Chỉ 5,1%', 'Việt Nam', 'V-NAMHS 2022'],
        ['VTN có rối nhiễu TL dùng dịch vụ', 'Chỉ 2%', 'Philippines', 'Mallari 2025'],
        ['Bác sĩ tâm thần/100.000 dân', '0,2 (Myanmar) \u2013 4,3 (Singapore)', 'ASEAN', 'GBD 2025'],
        ['Tỷ lệ NC dùng chẩn đoán chuẩn', 'Rất thấp', 'Nam Á', 'Mudunna 2025'],
    ],
    widths=[4.0, 3.0, 2.0, 2.5]
)
doc.add_paragraph()
p('Phân tích liên bài:', bold=True)
p('a) Phụ huynh là "gatekeeper" chính — nhưng chỉ 5,1% phụ huynh VN nhận ra con cần giúp đỡ (V-NAMHS). Đây là nút thắt lớn nhất: nếu phụ huynh không nhận ra, VTN không được tiếp cận dịch vụ dù dịch vụ có sẵn.')
p('b) So sánh VN (8,4%) vs Philippines (2%): Cả hai đều rất thấp, nhưng VN cao gấp 4 lần Philippines. Sự khác biệt có thể do VN có hệ thống y tế tốt hơn hoặc do khác biệt phương pháp đo.')
p('c) Thiếu nhân lực trầm trọng: Tỷ lệ bác sĩ tâm thần/100.000 dân tại ASEAN (0,2-4,3) thấp hơn rất nhiều so với trung bình châu Âu (9,7) (GBD 2025).')
p('Gap: Cần NC về rào cản tiếp cận dịch vụ SKTT ở VTN VN: kỳ thị, nhận thức, tài chính, địa lý (Jorm, 2012; Thornicroft và cộng sự, 2016).', italic=True)

doc.add_page_break()

# ===== 9. TỔNG HỢP GAP =====
h('9. Tổng hợp khoảng trống nghiên cứu (Research Gaps) — Ưu tiên cho Việt Nam')
tbl(
    ['#', 'Khoảng trống', 'Bằng chứng từ bài nào', 'Mức ưu tiên'],
    [
        ['1', 'RCT can thiệp SKTT tại trường học VN', 'Zhameden 2025: 0 RCT từ VN', 'RẤT CAO'],
        ['2', 'So sánh sàng lọc (GAD-7) vs chẩn đoán (DISC-5) trên CÙNG MẪU', 'Hoa 2024 vs V-NAMHS: chênh 17-24 lần', 'RẤT CAO'],
        ['3', 'NC dọc theo dõi CÙNG NHÓM HS trước-trong-sau COVID', 'Hoàng TH 2025, Puyat 2025: chỉ so 2 mẫu khác nhau', 'CAO'],
        ['4', 'Đo lường giấc ngủ \u2014 lo âu với AOR tại VN', 'Zhu 2025 (TQ): AOR=13,7 cho <5h', 'CAO'],
        ['5', 'RCT đào tạo chăm sóc cảm xúc cho cha mẹ/nhân viên SSF', 'Pham 2024: \u03b2=\u22120,40 nhưng cắt ngang', 'CAO'],
        ['6', 'NC SKTT HS DTTS mở rộng ngoài Lạng Sơn', 'Ngô Anh Vinh 2024: chỉ 1 tỉnh', 'TRUNG BÌNH'],
        ['7', 'Giải thích ngoại lệ V-NAMHS: tại sao không có khác biệt giới?', 'V-NAMHS vs tất cả NC khác', 'TRUNG BÌNH'],
        ['8', 'NC rào cản tiếp cận dịch vụ SKTT ở VTN VN', 'V-NAMHS: chỉ 8,4% tiếp cận', 'CAO'],
        ['9', 'So sánh ASEAN với cùng công cụ (VN vs Philippines vs Indonesia)', 'GBD 2025: tỷ lệ rất khác nhau', 'TRUNG BÌNH'],
        ['10', 'Chuẩn hóa DISC-5 cho văn hóa VN', 'V-NAMHS: DISC-5 chưa chuẩn hóa VN', 'CAO'],
    ],
    widths=[0.5, 4.5, 4.0, 2.0]
)

doc.add_paragraph()
p('Tài liệu tạo ngày 29/03/2026. Dựa trên 10 bài nghiên cứu đã dịch và kiểm tra (2022-2025).', italic=True, size=10)

fname = 'Tổng hợp liên bài báo - Lo âu HS - 29032026.docx'
doc.save(fname)
