"""DOC: Giai thich 'Long tu trong co cuong do tac dong tuong duong 85-89% voi
ap luc hoc tap' nghia la gi?
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Tu_trong_ngang_ALHT_85_89_phan_tram_giai_thich.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
GREEN = RGBColor(0x00, 0x70, 0x30)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('"LÒNG TỰ TRỌNG CÓ CƯỜNG ĐỘ TÁC ĐỘNG\nTƯƠNG ĐƯƠNG ~85–89% VỚI ÁP LỰC HỌC TẬP"\n— Nghĩa là thế nào? Trả lời thầy 09/05/2026 —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue(
    'Hôm trước bạn giải thích lòng tự trọng có cường độ tác '
    'động tương đương (~85–89%) với áp lực học tập nghĩa là '
    'thế nào?', italic=True
)

# Trả lời ngắn
H('Trả lời NGẮN GỌN', level=2, color=NAVY)
para(
    'Nghĩa là: nếu tăng áp lực học tập 1 đơn vị làm lo âu tăng '
    '1,0 đơn vị, thì tăng lòng tự trọng 1 đơn vị làm lo âu '
    'GIẢM khoảng 0,85–0,89 đơn vị. Hay nói cách khác: tự '
    'trọng có "lực" bảo vệ ngang bằng khoảng 85–89% "lực" '
    'gây lo âu của áp lực học tập.', bold=True
)

# I. Cách tính tỷ số
H('I. Cách tính tỷ số 85–89%', level=2, color=NAVY)
para('Em đã so sánh hệ số β chuẩn hóa của hai yếu tố trong chương 3 luận án:', bold=True)
add_table(
    ['Đường ảnh hưởng', 'β ALHT', 'β TTr', 'Tỷ số |β TTr| / |β ALHT|'],
    [
        ['→ RLLATQ (lan tỏa)', '0,510', '−0,455', '0,455 / 0,510 = 0,892 ≈ 89%'],
        ['→ RLLAXH (xã hội)', '0,490', '−0,415', '0,415 / 0,490 = 0,847 ≈ 85%'],
    ]
)
para('')
para('Tỷ số 85–89% nằm trong khoảng giữa hai con số trên.', bold=True)

# II. Hiểu β chuẩn hóa
H('II. Hiểu β chuẩn hóa là gì', level=2, color=NAVY)
para(
    'β (beta) chuẩn hóa cho biết: khi biến độc lập TĂNG 1 độ '
    'lệch chuẩn (SD), biến phụ thuộc THAY ĐỔI bao nhiêu SD.'
)
para('Ví dụ cụ thể:', bold=True)
para('• β ALHT → RLLATQ = 0,510 nghĩa là: khi ÁP LỰC HỌC TẬP tăng 1 SD, lo âu lan tỏa TĂNG 0,510 SD.')
para('• β TTr → RLLATQ = −0,455 nghĩa là: khi LÒNG TỰ TRỌNG tăng 1 SD, lo âu lan tỏa GIẢM 0,455 SD.')
para('')
para(
    'β chuẩn hóa CHO PHÉP SO SÁNH cường độ tác động giữa các '
    'biến vì đã loại bỏ ảnh hưởng của đơn vị đo (điểm, %, '
    'giờ...). Đây là lý do em dùng tỷ số |β TTr| / |β ALHT| '
    'để so sánh.'
)

# III. Ý nghĩa của dấu β
H('III. Vì sao dấu β khác nhau (DƯƠNG vs ÂM) nhưng vẫn so sánh được?', level=2, color=NAVY)
para('Dấu chỉ HƯỚNG, độ lớn |β| chỉ CƯỜNG ĐỘ:', bold=True)
add_table(
    ['Yếu tố', 'β', 'Hướng', 'Cường độ |β|'],
    [
        ['ALHT (áp lực học tập)', '+0,510', 'DƯƠNG → tăng lo âu (yếu tố nguy cơ)', '0,510'],
        ['TTr (lòng tự trọng)', '−0,455', 'ÂM → giảm lo âu (yếu tố bảo vệ)', '0,455'],
    ]
)
para('')
para(
    'Khi so sánh CƯỜNG ĐỘ giữa hai yếu tố, ta lấy GIÁ TRỊ TUYỆT '
    'ĐỐI |β|. Dấu DƯƠNG hay ÂM không ảnh hưởng đến cường độ — '
    'chỉ cho biết hướng tác động.'
)
para('Ví dụ minh họa:', bold=True)
para('• Một lực đẩy +10N và một lực kéo −10N có HƯỚNG ngược nhau, nhưng CƯỜNG ĐỘ bằng nhau (10N).')
para('• Tương tự, ALHT (+0,510) đẩy lo âu lên, TTr (−0,455) kéo lo âu xuống. Cường độ |0,455|/|0,510| = 89%.')

# IV. Diễn giải thực tế
H('IV. Diễn giải thực tế', level=2, color=NAVY)
para(
    'Hình dung trực quan: tưởng tượng áp lực học tập như một '
    '"lực kéo" làm tăng lo âu, lòng tự trọng như "lực kéo '
    'ngược" làm giảm lo âu.'
)
add_table(
    ['Tình huống giả định', 'Lực kéo lên (ALHT)', 'Lực kéo xuống (TTr)', 'Cân bằng?'],
    [
        ['Học sinh có ALHT cao + TTr THẤP', 'MẠNH (1,0×)', 'YẾU (0,3×)', 'Lo âu RẤT TĂNG'],
        ['Học sinh có ALHT cao + TTr CAO', 'MẠNH (1,0×)', 'MẠNH (0,89×)', 'Lo âu GẦN cân bằng'],
        ['Học sinh có ALHT thấp + TTr cao', 'YẾU (0,3×)', 'MẠNH (0,89×)', 'Lo âu RẤT GIẢM'],
    ]
)
para('')
para(
    'Nói cách khác: học sinh có TỰ TRỌNG CAO có thể PHẦN LỚN "vô '
    'hiệu hóa" tác động tiêu cực của áp lực học tập (89% cho '
    'lo âu lan tỏa, 85% cho lo âu xã hội).', bold=True
)

# V. Ý nghĩa cho can thiệp
H('V. Ý nghĩa cho can thiệp', level=2, color=NAVY)
para('Đây là phát hiện QUAN TRỌNG cho thiết kế chương trình tập huấn:', bold=True)
para(
    'Trong y văn quốc tế thường có giả định: "Yếu tố nguy cơ '
    'mạnh hơn yếu tố bảo vệ — phải ưu tiên giảm nguy cơ". Phát '
    'hiện chương 3 luận án TRÁI với giả định này — tự trọng có '
    'cường độ NGANG BẰNG 85–89% áp lực học tập.'
)
para('Hàm ý:', bold=True)
para('• Có thể can thiệp BẰNG hai con đường:')
para('   — (a) GIẢM áp lực học tập (giảm nguy cơ).')
para('   — (b) TĂNG lòng tự trọng (tăng bảo vệ).')
para('• Cả hai cách đều có hiệu quả TƯƠNG ĐƯƠNG về mặt cường độ.')
para('• Trong thực tế, khó GIẢM áp lực học tập (do hệ thống giáo dục, kỳ vọng cha mẹ) — nhưng TĂNG TỰ TRỌNG khả thi hơn (qua chương trình tập huấn).')
para('')
para(
    'Phù hợp với khung "ordinary magic" của Masten (2014): '
    'lòng tự trọng là yếu tố nội tại bảo vệ trẻ em trước '
    'nghịch cảnh — đa số trẻ có khả năng tăng tự trọng khi '
    'được hỗ trợ phù hợp.'
)

# VI. Cảnh báo khi áp dụng
H('VI. Cảnh báo khi áp dụng', level=2, color=NAVY)
para('Bốn lưu ý quan trọng:', bold=True)
para(
    '(1) Tỷ số 85–89% chỉ áp dụng cho HAI dạng RLLA cụ thể (lan '
    'tỏa và xã hội). Đối với LO ÂU CHIA LY (RLLAC), β TTr = '
    '−0,087 — yếu hơn RẤT NHIỀU so với β ALHT (0,253) — chỉ '
    '~34%.'
)
para(
    '(2) Tỷ số này đo CƯỜNG ĐỘ trên thang chuẩn hóa SD, KHÔNG '
    'phải tỷ số "giải thích phương sai" R². R² của TTr riêng '
    '= 0,209 (Bảng 3.32) — thấp hơn R² của ALHT riêng = 0,284 '
    '(Bảng 3.24). Tỷ số 85–89% là về β, không phải về R².'
)
para(
    '(3) Mô hình SEM RIÊNG cho TTr → RLLA — chưa kiểm soát các '
    'yếu tố khác (như ALHT). Khi tích hợp YTBV với YTNC (Bảng '
    '3.38), |β YTBV| GIẢM còn 0,220 — vì có TƯƠNG QUAN giữa '
    'TTr và ALHT.'
)
para(
    '(4) Phát hiện "85–89%" CHƯA được kiểm chứng trong nghiên '
    'cứu can thiệp (RCT). Cần thử nghiệm thực tế: can thiệp '
    'tăng tự trọng có giảm lo âu tương đương giảm áp lực không?'
)

# VII. Tóm gọn
H('VII. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue('Cách tính:', bold=True)
blue('• |β TTr → RLLATQ| / |β ALHT → RLLATQ| = 0,455 / 0,510 = 0,892 ≈ 89%')
blue('• |β TTr → RLLAXH| / |β ALHT → RLLAXH| = 0,415 / 0,490 = 0,847 ≈ 85%')
blue('')
blue('Diễn giải:', bold=True)
blue('Khi áp lực học tập tăng 1 SD làm lo âu tăng 1,0 SD, thì '
     'tăng lòng tự trọng 1 SD làm lo âu GIẢM khoảng 0,85–0,89 SD. '
     'Tự trọng có "lực" bảo vệ ngang bằng 85–89% "lực" gây lo '
     'âu của áp lực học tập.')
blue('')
blue('Hàm ý cho can thiệp:', bold=True)
blue('Có thể giảm lo âu BẰNG hai cách: (a) giảm áp lực học tập, '
     '(b) tăng lòng tự trọng. Cả hai có hiệu quả TƯƠNG ĐƯƠNG. '
     'Trong thực tế, TĂNG TỰ TRỌNG khả thi hơn qua chương trình '
     'tập huấn — vì khó giảm áp lực học tập do hệ thống giáo '
     'dục và kỳ vọng cha mẹ.')
blue('')
blue('Lưu ý:', bold=True)
blue('Tỷ số 85–89% CHỈ áp dụng cho lo âu lan tỏa và xã hội. '
     'Đối với lo âu chia ly, tự trọng yếu hơn nhiều (chỉ ~34% '
     'cường độ áp lực học tập).')

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
