"""Canonical hóa QT065 (Lochman/Bradshaw 2025 EACP) + QT066 (Murphy 2024 Peer Support).
Metadata-only — KHÔNG có PDF (chờ thầy tải về). Build summary docx Jenkins template.
"""
import sys, io, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
TT_DIR = BASE / 'Tom-tat-tung-bai'

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)

PAPERS = [
    {
        'cid': 'QT065',
        'descriptor': 'Bradshaw_Lochman_EACP_RCT_JSchPsy_2025',
        'pdf_dst': 'QT065_Bradshaw_Lochman_EACP_RCT_JSchPsy_2025.pdf',
        'folder': 'The-gioi_Au-My-Uc',
        'title_vn': 'Thử nghiệm RCT Chương trình "Sức mạnh Ứng phó" (Early Adolescent Coping Power): Hiệu quả lên các vấn đề cảm xúc và hành vi của học sinh THCS',
        'title_en': 'Randomized Controlled Trial of the Early Adolescent Coping Power Program: Effects on Emotional and Behavioral Problems in Middle Schoolers',
        'authors': 'Bradshaw CP, McDaniel H, Pas ET, Debnam KJ, Bottiani JH, Powell N, Ialongo NS, Morgan-Lopez A, Lochman JE',
        'lead_author_email': 'cpb6n@virginia.edu (Catherine P. Bradshaw, UVA)',
        'affiliation': 'UVA + USC + Johns Hopkins + UA + RTI International (multi-site)',
        'journal': 'Journal of School Psychology, vol. 110, June 2025, art. 101437',
        'year': 2025,
        'doi': '10.1016/j.jsp.2025.101437',
        'pmc': 'PMC12163234 (embargo đến 2026-06-01)',
        'pubmed': 'https://pubmed.ncbi.nlm.nih.gov/40506167/',
        'sample': '40 trường THCS Mỹ, n=709 HS lớp 7 (sàng lọc giáo viên cho hành vi gây hấn). 69,8% gốc Phi (African American), 59,4% nam.',
        'location': 'Hoa Kỳ (multi-site)',
        'tool': 'BASC (Behavior Assessment System for Children) — teacher-rated + self-rated',
        'design': 'Cluster RCT — EACP vs treatment-as-usual; 1 năm intervention + 3 năm follow-up',
        'method': 'RCT cluster school-level',
        'outcomes': [
            'Cấu trúc EACP: 25 buổi học sinh + 16 buổi cha mẹ + thành phần giáo viên (giảm từ bản gốc 34 buổi)',
            'Skills: quản lý cơn giận-lo âu (self-statements + thư giãn) + giải quyết vấn đề xã hội + tương tác bạn bè + kỹ năng học tập + từ chối chất gây nghiện',
            'Teacher-rated: EACP giảm externalizing problems theo thời gian',
            'Đặc biệt: NỮ trong EACP giảm school problems có ý nghĩa thống kê',
            'Self-rated: NỮ trong EACP có personal adjustment outcomes tốt hơn (moderated effect)',
            'Hiệu ứng phòng ngừa cho HS lớp 7; promotive effects cho NỮ',
        ],
        'strengths': [
            'RCT cluster đa trung tâm 40 trường — bằng chứng cao',
            'Mẫu lớn (n=709) trong dân số khó tiếp cận (HS gây hấn 69,8% African American)',
            'Tài trợ NIMH — methodology vững',
            'J School Psychology Q1 — peer-reviewed top',
            'Theo dõi đến 3 năm post-intervention',
            'BASC — thang đo chuẩn lâm sàng',
        ],
        'limitations': [
            '⚠ Đối tượng GỐC là HS gây hấn (aggressive behavior), KHÔNG phải HS lo âu thuần',
            'Outcome chính là externalizing (hành vi), không phải lo âu (internalizing)',
            'Chủ yếu dân Mỹ gốc Phi (69,8%) — generalize cho VN cần adapt văn hoá',
            'Chỉ HIỆU QUẢ rõ với NỮ — moderated effect',
            'Không đo trực tiếp triệu chứng lo âu (chỉ qua personal adjustment)',
        ],
        'critique_short': 'EACP là chương trình nổi tiếng (Lochman) với 25+ năm phát triển. Bài 2025 này test phiên bản RÚT GỌN (25 buổi/1 năm) trên HS THCS lớp 7. PHÁT HIỆN MỚI: hiệu quả ĐẶC BIỆT với NỮ — gợi ý EACP phù hợp cho HS NỮ với cả vấn đề externalizing + internalizing. CẢNH BÁO: đối tượng gốc là HS gây hấn (không phải lo âu thuần), nên áp dụng cho HS lo âu VN cần modify curriculum.',
        'future': 'Adapt EACP cho HS THCS VN với focus trên LO ÂU (modify Module quản lý cơn giận thành quản lý lo âu). Test cluster RCT 5-10 trường VN, n>200 HS. Đo bằng GAD-7 + DASS-21 thay vì BASC.',
        'rating': 5,
        'notes_for_rag': 'EACP RCT 40 trường, n=709 HS lớp 7 Mỹ. 25 buổi HS + 16 buổi cha mẹ. Hiệu quả đặc biệt với NỮ. Đối tượng gốc là HS gây hấn.',
        'has_pdf': False,
    },
    {
        'cid': 'QT066',
        'descriptor': 'Murphy_PeerSupport_ScopingReview_JCommPsych_2024',
        'pdf_dst': 'QT066_Murphy_PeerSupport_ScopingReview_JCommPsych_2024.pdf',
        'folder': 'The-gioi_Au-My-Uc',
        'title_vn': 'Tổng quan phạm vi hệ thống về can thiệp hỗ trợ ngang hàng (peer support) trong chăm sóc sức khoẻ tâm thần thanh niên',
        'title_en': 'A systematic scoping review of peer support interventions in integrated primary youth mental health care',
        'authors': 'Murphy R, Huggard L, Fitzgerald A, Hennessy E, Booth A',
        'lead_author_email': 'rachel.murphy@ucd.ie (Rachel Murphy, University College Dublin)',
        'affiliation': 'University College Dublin (Ireland) + Jigsaw national youth MH service',
        'journal': 'Journal of Community Psychology, vol. 52(1):154-180 (2024)',
        'year': 2024,
        'doi': '10.1002/jcop.23090',
        'pubmed': 'https://pubmed.ncbi.nlm.nih.gov/37740958/',
        'sample': '15 nghiên cứu gốc, 13 interventions peer support trong primary youth MH care',
        'location': 'Quốc tế — chủ yếu Ireland (Jigsaw), UK, Mỹ, Úc',
        'tool': 'GAD-7, DASS-21, recovery-related outcomes, self-efficacy scales',
        'design': 'Scoping review (Arksey & O\'Malley framework)',
        'method': 'Scoping review',
        'outcomes': [
            'Định nghĩa peer support: "hỗ trợ xã hội + cảm xúc giữa cá nhân có TRẢI NGHIỆM CHUNG về khó khăn MH"',
            '8/13 interventions nhắm youth có khó khăn MH NHẸ-VỪA (trầm cảm, lo âu, distress)',
            'Phân bố tuổi: 9/13 nhắm 16-25 (thanh niên); 1 nhắm 13-19; 1 nhắm 11-18 (gần HS THCS); 1 nhắm 12-25',
            'Phát hiện: peer support có TIỀM NĂNG cải thiện outcome PHỤC HỒI (recovery)',
            'Rào cản: lo ngại bảo mật + peer worker thiếu tự tin',
            'Yếu tố thúc đẩy: nhân viên chuyên môn HỖ TRỢ + vai trò peer được ĐỊNH NGHĨA RÕ',
            'Khuyến nghị: nghiên cứu sau cần báo cáo CHI TIẾT HƠN nội dung intervention',
        ],
        'strengths': [
            'Scoping review chuẩn methodology (Arksey & O\'Malley)',
            'Bài đầu tiên tổng hợp peer support trong primary youth MH care',
            'Wiley J Community Psychology — peer-reviewed',
            'Định nghĩa peer support rõ ràng',
            'Phân tích cả barrier + facilitator',
        ],
        'limitations': [
            '⚠ Đa số 9/13 interventions cho 16-25 (thanh niên), CHỈ 2/13 có HS THCS-THPT',
            'Scoping (không phải systematic + MA) — không đo effect size',
            'Nội dung intervention thiếu chi tiết → khó nhân rộng',
            'Bias bối cảnh: chủ yếu Ireland (Jigsaw) — không đại diện toàn cầu',
            'Không có RCT nào trong 13 interventions test riêng peer support cho HS THCS',
        ],
        'critique_short': 'Bài tổng quan QUÝ về peer support — nhưng đối tượng gốc đa số THANH NIÊN 16-25. Để áp dụng cho HS THCS-THPT VN, cần adapt mạnh: (1) phù hợp tuổi 11-15; (2) bảo mật + parental consent; (3) đào tạo peer worker tuổi gần. Mô hình peer support khả thi cho VN vì: HS VN ưu thích NON-PROFESSIONAL trước (gia đình, bạn bè) — peer support đáp ứng đúng preference này.',
        'future': 'Tổng quan systematic + MA định lượng cho peer support ở HS THCS-THPT. RCT VN: 2-3 trường thí điểm peer support (HS lớp 11 đào tạo làm peer leader cho HS lớp 7-8) + đo GAD-7 + DASS-21.',
        'rating': 4,
        'notes_for_rag': 'Murphy 2024 scoping review 15 NC peer support youth MH. Đa số 16-25 tuổi. Khuyến nghị peer worker được hỗ trợ + vai trò rõ. Thích hợp adapt cho VN vì HS VN ưu thích non-professional.',
        'has_pdf': False,
    },
]

def build_summary(paper, out_path):
    d = Document()
    style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
    DARK = RGBColor(31,73,125); GREEN = RGBColor(54,95,44); RED = RGBColor(192,0,0); ORANGE = RGBColor(191,97,14)

    t = d.add_heading(f"{paper['cid']} — Tóm tắt bài nghiên cứu", level=0)
    for r in t.runs: r.font.color.rgb = DARK

    if not paper['has_pdf']:
        warn = d.add_table(rows=1, cols=1); warn.style = 'Table Grid'
        wc = warn.rows[0].cells[0]; shade(wc, 'FFF2CC')
        wp = wc.paragraphs[0]
        wr = wp.add_run('⚠ Lưu ý: '); wr.bold = True; wr.font.color.rgb = ORANGE
        wp.add_run(f"Tóm tắt này dựa trên metadata (PubMed + abstract + thông tin cào web), CHƯA có PDF gốc. "
                   f"Thầy tải PDF từ: DOI {paper['doi']} hoặc PubMed {paper['pubmed']}. "
                   f"Sau khi có PDF, em sẽ verify số liệu chi tiết.")
        d.add_paragraph()

    meta = d.add_table(rows=10, cols=2); meta.style = 'Table Grid'
    mrows = [
        ('Tên (VN)', paper['title_vn']),
        ('Tên (EN)', paper['title_en']),
        ('Tác giả', paper['authors']),
        ('Lead author email', paper['lead_author_email']),
        ('Đơn vị', paper['affiliation']),
        ('Tạp chí', paper['journal']),
        ('Năm', str(paper['year'])),
        ('DOI', paper['doi']),
        ('PubMed', paper.get('pubmed', 'N/A')),
        ('Mẫu + công cụ', f"{paper['sample']} | Công cụ: {paper['tool']}"),
    ]
    for i, (k, v) in enumerate(mrows):
        c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
        p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
        meta.rows[i].cells[1].text = v
    d.add_paragraph()

    for sec, items, color in [
        ('Phương pháp', [paper['design'], f"Mẫu: {paper['sample']}", f"Công cụ: {paper['tool']}"], DARK),
        ('Kết quả nghiên cứu', paper['outcomes'], DARK),
        ('Điểm mạnh', paper['strengths'], GREEN),
        ('Hạn chế', paper['limitations'], RED),
    ]:
        h = d.add_heading(sec, level=1)
        for r in h.runs: r.font.color.rgb = color
        for it in items: d.add_paragraph('• ' + it)

    h = d.add_heading('Phản biện ngắn', level=1)
    for r in h.runs: r.font.color.rgb = RED
    p = d.add_paragraph(); rr = p.add_run(paper['critique_short']); rr.font.color.rgb = RED

    h = d.add_heading('Hướng nghiên cứu tiếp theo', level=1)
    for r in h.runs: r.font.color.rgb = DARK
    d.add_paragraph(paper['future'])

    stars = '⭐' * paper['rating']
    h = d.add_heading(f"Đánh giá: {stars} ({paper['rating']}/5)", level=1)
    for r in h.runs: r.font.color.rgb = ORANGE

    if paper.get('notes_for_rag'):
        d.add_paragraph()
        p = d.add_paragraph(); r = p.add_run('Key fact cho RAG: '); r.bold = True; r.font.size = Pt(10)
        p.add_run(paper['notes_for_rag']).font.size = Pt(10)

    d.save(out_path)

# Execute
log = []
with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

for paper in PAPERS:
    cid = paper['cid']
    tt_name = f"{cid}_{paper['descriptor']}.docx"
    tt_path = TT_DIR / tt_name
    build_summary(paper, tt_path)
    log.append(f"SUMMARY {cid}: {tt_name} ({tt_path.stat().st_size//1024} KB)")

    idx[cid] = {
        'id': cid,
        'descriptor': paper['descriptor'],
        'summary': tt_name,
        'translation': None,
        'pdf': None if not paper['has_pdf'] else paper['pdf_dst'],
        'pdf_folder': paper['folder'] if paper['has_pdf'] else None,
        'doi': paper['doi'],
        'pubmed': paper.get('pubmed'),
        'pmc': paper.get('pmc'),
        'lead_author_email': paper.get('lead_author_email'),
        'cao_date': '2026-04-29',
        'scope': 'main',
        'method': paper.get('method', ''),
        'has_pdf': paper['has_pdf'],
        'pdf_pending_filename': paper['pdf_dst'],  # tên kỳ vọng khi tải về
    }
    log.append(f"INDEX {cid}: added (no PDF — pending download)")

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

n_vn = sum(1 for k in idx if k.startswith('VN'))
n_qt = sum(1 for k in idx if k.startswith('QT'))
log.append(f'\nTotal canonical: {len(idx)} ({n_vn} VN + {n_qt} QT)')
log.append('\nCHỜ THẦY TẢI PDF:')
for paper in PAPERS:
    log.append('  - ' + paper['cid'] + ': ' + paper['pdf_dst'])
    log.append('    DOI: ' + paper['doi'])
    log.append('    PubMed: ' + paper.get('pubmed', ''))

for l in log: print(l)
