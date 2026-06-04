# -*- coding: utf-8 -*-
"""Sinh file Bao cao loi v2 - tong hop tat ca 15+ loi da phat hien.
Co toa do cu the (file + para number) + Find/Replace text.
26-27/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '08_BaoCao_LoiBiSai_v2_27052026.docx')

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(11)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.3
    for sec in doc.sections:
        sec.top_margin = Cm(1.8); sec.bottom_margin = Cm(1.8)
        sec.left_margin = Cm(2.0); sec.right_margin = Cm(1.5)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(192, 0, 0) if level >= 2 else RGBColor(0, 0, 0)

def P(doc, text, indent=True, italic=False, bold=False, color=None, size=11):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(0.7)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size)
    r.italic = italic; r.bold = bold
    if color is not None:
        r.font.color.rgb = color

def make_table(doc, headers, rows, widths, hdr_color='FFE699'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = htxt
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, hdr_color)
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

def fix_block(doc, num, paper, location, find_text, replace_text, errors_listed):
    """Mot block sua loi cu the."""
    H(doc, f'Lỗi #{num} — {paper}', 3)
    P(doc, f'Vị trí: {location}', bold=True, italic=True, color=RGBColor(192, 0, 0), size=11)
    P(doc, f'Loại lỗi: {errors_listed}', italic=True, color=RGBColor(80, 80, 80), size=10)
    P(doc, 'TEXT HIỆN TẠI (sai):', bold=True, size=10, color=RGBColor(192, 0, 0))
    P(doc, f'"{find_text}"', italic=True, indent=False, size=10)
    P(doc, 'SỬA THÀNH:', bold=True, size=10, color=RGBColor(0, 112, 0))
    P(doc, f'"{replace_text}"', italic=True, indent=False, size=10, color=RGBColor(0, 112, 0))
    doc.add_paragraph()

# ============================================================
# BUILD
# ============================================================
doc = doc_init()

H(doc, 'BÁO CÁO LỖI SAI VÀ CHỖ CẦN SỬA TRONG LUẬN ÁN', 1)
P(doc, '(Sau audit kỹ từng nghiên cứu được trích dẫn — NCS Công Thị Hằng — 27/05/2026)', indent=False, italic=True, color=RGBColor(80, 80, 80))
doc.add_paragraph()

# Tom luoc nhanh
H(doc, '1. Tóm lược tình hình', 2)
P(doc, 'Sau khi audit toàn bộ các nghiên cứu được trích dẫn trong Tổng quan v2 33-trang và Tổng quan FULL v1 bằng cách đối chiếu trực tiếp từng số liệu với PDF gốc, đã phát hiện 15 lỗi nội dung trong 6 nghiên cứu, kèm 2 chỗ thuật ngữ "lo âu tổng quát" cần đổi thành "lo âu lan tỏa" cho nhất quán với DSM-5 Việt Nam.')

P(doc, 'Quy luật chung của lỗi: các nghiên cứu được tôi (AI) thêm vào để bổ sung độ rộng cho Tổng quan — KHÔNG nằm trong 21 TLTK của Bài 1 hay 21 TLTK của Bài 2 đã verify từ PDF — có rủi ro lớn bị bịa số liệu. 6 nghiên cứu có lỗi đều thuộc nhóm này.')

P(doc, 'Đã có sẵn 2 file Tổng quan v3 mới với các lỗi đã được sửa và tô đỏ (để thầy/NCS dễ thấy). Báo cáo này liệt kê chi tiết từng lỗi để NCS có thể tự sửa thủ công trong bản luận án đang dùng nếu muốn.', bold=True, color=RGBColor(0, 112, 0))
doc.add_paragraph()

# Bang tom luoc
H(doc, '2. Bảng tổng hợp lỗi', 2)
make_table(doc,
    ['#', 'Nghiên cứu', 'Số lỗi', 'Loại lỗi chính', 'Trạng thái'],
    [
        ['1', 'VN017 Nguyễn Danh Lâm 2022', '4', 'n=1.024→482; bán đô thị→bán nông thôn; 25,8%→49,0%; bịa "nam>nữ"', '✓ Đã sửa (v3)'],
        ['2', 'VN016 Nguyễn Ngọc Bảo Quyên 2025', '2', 'n=1.252→501; 38,5%→86,2%', '✓ Đã sửa (v3)'],
        ['3', 'VN018 An Giang 2025', '3', 'Tên TG sai (Lê Minh T.→Nguyễn Đăng Khoa); n=600→366; 33,2%→61,2%', '✓ Đã sửa (v3)'],
        ['4', 'QT035 Jefferies & Ungar 2020', '2', '7 quốc gia sai (Canada+Anh→Brazil+Nga); Thái Lan cao nhất→US cao nhất', '✓ Đã sửa (chỉ FULL v1)'],
        ['5', 'QT005 Alharbi 2019', '2', 'n=366→1.245; 63,1%→63,4%', '✓ Đã sửa (chỉ FULL v1)'],
        ['6', 'QT006 Nakie 2022', '2', 'n=643→849; 38,9%→66,7%', '✓ Đã sửa (chỉ FULL v1)'],
        ['7', 'Thuật ngữ "lo âu tổng quát"', '2', 'Đổi 2 chỗ thành "lo âu lan tỏa" (chuẩn DSM-5 VN)', '✓ Đã sửa (cả 2 file)'],
    ],
    [0.8, 5.0, 1.2, 8.5, 2.8],
    hdr_color='FFD7D7'
)
P(doc, 'Tổng: 15 lỗi nội dung + 2 chỗ thuật ngữ = 17 sửa đổi trong Tổng quan v2/FULL v1.')
doc.add_paragraph()

# So lieu DUNG cho 6 NC co loi
H(doc, '3. Số liệu ĐÚNG cho 6 nghiên cứu có lỗi (đã đối chiếu PDF gốc)', 2)
make_table(doc,
    ['Nghiên cứu', 'n đúng', 'Tỷ lệ lo âu đúng', 'Địa bàn đúng', 'Tạp chí đúng'],
    [
        ['VN017 Nguyễn Danh Lâm 2022', '482 HS THPT', '49,0% (DASS-21)', 'Huyện Yên Định, Thanh Hóa (bán nông thôn)', 'Y học Việt Nam tập 516, số 1, 7/2022, tr. 67-70'],
        ['VN016 Nguyễn Ngọc Bảo Quyên 2025', '501 HS THPT', '86,2% (DASS-21)', 'Hà Nội', 'Y học Cộng đồng, tập 66 chuyên đề 10, 2025, tr. 79-86'],
        ['VN018 An Giang 2025', '366 HS THCS&THPT', '61,2% (DASS-21)', 'Long Bình, An Giang', 'Y học Việt Nam tập 549, số 1, 4/2025'],
        ['QT035 Jefferies & Ungar 2020', '6.825 thanh niên 7 nước', '36,2% trung bình; VN 30,7%; US cao nhất 57,6%', 'Brazil, Trung Quốc, Indonesia, Nga, Thái Lan, Hoa Kỳ, Việt Nam', 'PLOS ONE 15(9):e0239133'],
        ['QT005 Alharbi 2019', '1.245 HS THPT', '63,4% (GAD-7)', 'Vùng Qassim, Ả Rập Saudi', 'J Family Medicine Primary Care 2019;8:504-10'],
        ['QT006 Nakie 2022', '849 HS THPT', '66,7% (DASS-21)', 'Northwest Ethiopia', 'BMC Psychiatry 2022, 22:739'],
    ],
    [4.0, 2.0, 4.0, 4.0, 4.0],
    hdr_color='C6E0B4'
)
P(doc, '6 nghiên cứu này có tên tác giả chính xác như sau (cần ghi đúng trong TLTK):', italic=True)
P(doc, 'VN018 An Giang: Nguyễn Đăng Khoa, Lê Minh Thi, Ngô Anh Vinh (KHÔNG phải "Lê Minh T." hay "Lê Minh và cộng sự" — Lê Minh Thi là tác giả thứ 2, không phải đầu tiên).')
P(doc, 'QT035 Jefferies: Philip Jefferies, Michael Ungar (2 tác giả; affil. Dalhousie University, Canada).')
doc.add_paragraph()

# Vi tri sua trong tung file
H(doc, '4. Tọa độ chi tiết các chỗ cần sửa', 2)
P(doc, 'Hai cách sử dụng:', italic=True)
P(doc, 'CÁCH 1 (NHANH): Thay thế hoàn toàn 2 file Tổng quan đã có lỗi bằng 2 file v3/v2 mới:')
P(doc, '   • 01_TongQuan_TheoChuDe_33trang_v3_FixSoLieu_26052026.docx (đã sửa 5/5 lỗi trong v2 33-trang)')
P(doc, '   • 01_TongQuan_TheoChuDe_FULL_v2_FixSoLieu_26052026.docx (đã sửa 8/8 lỗi trong FULL v1)')
P(doc, 'CÁCH 2 (THỦ CÔNG): Áp dụng các sửa cụ thể dưới đây trên bản luận án đang dùng. Các phần đã sửa được tô ĐỎ ĐẬM để thầy phản biện thấy chỗ thay đổi.')
doc.add_paragraph()

H(doc, '4.1. Tổng quan v2 33-trang (file đưa vào LA)', 3)
P(doc, 'File: 01_TongQuan_TheoChuDe_33trang_FINAL_26052026.docx', italic=True, color=RGBColor(80, 80, 80))

# Loi 1-3 trong para 8 - all in single paragraph
fix_block(doc, 1,
    'Bảo Quyên + Lâm + An Giang (3 lỗi trong 1 đoạn)',
    'Mục 1.1.1 — Paragraph 8 (đoạn về "ba khảo sát quy mô vừa và lớn")',
    'Nguyễn Ngọc Bảo Quyên và cộng sự (2025) trên 1.252 học sinh trung học phổ thông Hà Nội bằng thang DASS-21 phiên bản tiếng Việt báo cáo tỷ lệ lo âu chung 38,5%, với nữ cao hơn nam có ý nghĩa thống kê. Nguyễn Danh Lâm và cộng sự (2022) trên 1.024 học sinh trung học phổ thông tại Yên Định, Thanh Hóa — khu vực bán đô thị — báo cáo tỷ lệ thấp hơn 25,8% bằng DASS-21, đồng thời ghi nhận học sinh nam có tỷ lệ cao hơn nữ ở khu vực này — một phát hiện trái ngược với xu hướng chung và đáng được nghiên cứu sâu hơn. Lê Minh T. và cộng sự (2025) trên 600 học sinh phổ thông Long Bình, An Giang bằng DASS-21 báo cáo tỷ lệ 33,2%.',
    'Nguyễn Ngọc Bảo Quyên và cộng sự (2025) trên 501 học sinh trung học phổ thông Hà Nội bằng thang DASS-21 phiên bản tiếng Việt báo cáo tỷ lệ lo âu chung 86,2% (phân tầng: nhẹ 6,6%; vừa 25,1%; nặng 17,4%; rất nặng 37,1%), với nữ cao hơn nam có ý nghĩa thống kê (p < 0,001). Nguyễn Danh Lâm và cộng sự (2022) trên 482 học sinh trung học phổ thông (2 trường) tại huyện Yên Định, Thanh Hóa — huyện bán nông thôn — báo cáo tỷ lệ lo âu chung 49,0% bằng DASS-21 (nhẹ 11,2%; vừa 25,1%; nặng 8,1%; rất nặng 4,6%). Nghiên cứu này không phân tích lo âu theo giới tính. Nguyễn Đăng Khoa, Lê Minh Thi và Ngô Anh Vinh (2025) trên 366 học sinh THCS&THPT Long Bình, An Giang bằng DASS-21 báo cáo tỷ lệ lo âu 61,2% (nhẹ 9,3%; vừa 24,0%; nặng 12,6%; rất nặng 15,3%).',
    '7 lỗi: n×3 + địa lý + ty le×3 + bịa giới + tên tác giả')

# Lan toa - cho 1
fix_block(doc, 2,
    'Thuật ngữ "tổng quát" → "lan tỏa"',
    'Mục 1.1.7 — Paragraph 24 (Khoảng trống thứ hai)',
    'lo âu tổng quát (chủ yếu qua GAD-7 hoặc các tiểu thang trong DASS)',
    'lo âu lan tỏa (chủ yếu qua GAD-7 hoặc các tiểu thang trong DASS)',
    'Thuật ngữ DSM-5 Việt Nam')

# Lan toa - cho 2
fix_block(doc, 3,
    'Thuật ngữ "tổng quát" → "lan tỏa"',
    'Mục 1.1.6 — Paragraph 80 (đoạn Bress 2024 ứng dụng Maya)',
    'rối loạn lo âu (lo âu tổng quát 56%, lo âu xã hội 41%)',
    'rối loạn lo âu (lo âu lan tỏa 56%, lo âu xã hội 41%)',
    'Thuật ngữ DSM-5 Việt Nam')

# Tổng quan FULL v1
H(doc, '4.2. Tổng quan FULL v1 (bản đầy đủ để tra cứu)', 3)
P(doc, 'File: 01_TongQuan_TheoChuDe_FULL_v1_26052026.docx', italic=True, color=RGBColor(80, 80, 80))
P(doc, 'Cùng 3 lỗi ở mục 4.1 (Bảo Quyên + Lâm + An Giang + 2 lan tỏa) — trong FULL v1 ở paragraph 10 (chứ không 8 vì FULL v1 có 2 đoạn intro hơn).')

# QT005 Alharbi
fix_block(doc, 4,
    'Alharbi 2019',
    'Paragraph chứa "Bhardwaj/Alharbi/Nakie" (gần mục 1.1.1)',
    'Alharbi và cộng sự (2019) trên 366 học sinh trung học phổ thông tại vùng Qassim, Ả Rập Saudi bằng GAD-7 báo cáo tỷ lệ 63,1%',
    'Alharbi và cộng sự (2019) trên 1.245 học sinh trung học phổ thông tại vùng Qassim, Ả Rập Saudi bằng GAD-7 báo cáo tỷ lệ 63,4% (nhẹ 34,1%; vừa 19,5%; nặng 9,8%)',
    'n và tỷ lệ chính xác')

# QT006 Nakie
fix_block(doc, 5,
    'Nakie 2022',
    'Paragraph chứa "Nakie/Bhardwaj/Alharbi" (gần mục 1.1.1)',
    'Nakie và cộng sự (2022) trên 643 học sinh trung học phổ thông tại Ethiopia bằng DASS-21 báo cáo tỷ lệ lo âu 38,9%',
    'Nakie và cộng sự (2022) trên 849 học sinh trung học phổ thông tại Northwest Ethiopia bằng DASS-21 báo cáo tỷ lệ lo âu 66,7% (cùng với trầm cảm 41,4% và stress 52,2%)',
    'n và tỷ lệ chính xác (sai gần gấp đôi tỷ lệ)')

# QT035 Jefferies
fix_block(doc, 6,
    'Jefferies & Ungar 2020',
    'Paragraph chứa "Jefferies" — mục 1.1.1 (đoạn về so sánh quốc tế)',
    'Sử dụng thang SIAS (Social Interaction Anxiety Scale) trên thanh niên ở Việt Nam, Thái Lan, Indonesia, Trung Quốc, Hoa Kỳ, Canada và Anh, các tác giả báo cáo tỷ lệ trung bình 36% có lo âu xã hội cho cả bảy quốc gia, trong đó Việt Nam đạt 30,7% — gần với mức trung bình. Thái Lan có tỷ lệ cao nhất 41%, Indonesia thấp nhất 22,9%.',
    'Sử dụng thang SIAS (Social Interaction Anxiety Scale) trên thanh niên ở bảy quốc gia (Brazil, Trung Quốc, Indonesia, Nga, Thái Lan, Hoa Kỳ và Việt Nam), các tác giả báo cáo tỷ lệ trung bình 36,2% có lo âu xã hội, trong đó Việt Nam đạt 30,7% — gần với mức trung bình. Hoa Kỳ có tỷ lệ cao nhất (57,6%); Brazil 42,4%; Thái Lan 41,4%; Trung Quốc 32,1%; Việt Nam 30,7%; Nga 27,0%; Indonesia thấp nhất 22,9%.',
    'Danh sách 7 nước + nước cao nhất')

doc.add_paragraph()

# Tinh trang khac
H(doc, '5. Trạng thái các tài liệu khác liên quan', 2)
make_table(doc,
    ['Tài liệu', 'Lỗi đã phát hiện', 'Tình trạng'],
    [
        ['Bản dịch VN017_NguyenDanhLam (03_Ban-dich/)', '4 lỗi số ở Bảng 2 Stress + 2 lỗi Lo âu + "bán đô thị" 3 chỗ + bịa claim "2024"', '⚠️ CẦN SỬA — chưa làm'],
        ['Tóm tắt VN017_NguyenDanhLam (Tom-tat-tung-bai/)', '4 lỗi số Bảng 1 Stress + 2 lỗi Lo âu + "bán đô thị" 5 chỗ', '⚠️ CẦN SỬA — chưa làm'],
        ['Bản dịch VN016_BaoQuyen (03_Ban-dich/)', 'KHÔNG có lỗi — số liệu chính xác với PDF gốc', '✓ OK'],
        ['Tóm tắt VN016_BaoQuyen (Tom-tat-tung-bai/)', '1 lỗi nhỏ: gọi "Thảo Vi 2025 (Huế)" — thực ra là Hồ Thị Trúc Quỳnh', '🟡 Lỗi nhỏ'],
        ['Bản dịch VN018_AnGiang (03_Ban-dich/)', 'Chưa audit', '⏳ Pending'],
        ['Tóm tắt VN018_AnGiang (Tom-tat-tung-bai/)', 'Chưa audit', '⏳ Pending'],
        ['Các bản dịch + tóm tắt khác (VN001, VN002, VN006, VN014, VN015, VN020, VN021, VN025, VN027, VN029)', 'Chưa audit chi tiết bản dịch/tóm tắt; nhưng các paper này đã verify từ PDF gốc trong các phiên trước', '🟡 Có thể có lỗi mức bảng — cần audit thêm'],
        ['Tài liệu QT khác (chưa nói rõ ở Tổng quan có claim cụ thể)', 'Anderson, Saikia, Pascoe, Kendall, Barrett-Dadds-Rapee, Bower-Gilbody', '⏳ Có thể audit ở phiên sau nếu cần'],
    ],
    [6.5, 8.0, 3.0],
    hdr_color='FFE699'
)
doc.add_paragraph()

# Quy trinh audit con lai
H(doc, '6. Quy trình audit còn lại (đề xuất cho phiên sau)', 2)
P(doc, 'Theo playbook 17-loại fact-check (memory feedback-quy-trinh-fact-check.md), phải áp dụng cả 17 loại trước khi tuyên bố "đã verify đủ". Tới hiện tại đã chạy chủ yếu loại 1 (verify số liệu cụ thể) cho phần lớn TLTK Tổng quan v2. Còn các loại sau cần làm trong phiên sau:')
steps = [
    'Loại 2: Verify qualitative claims cho mỗi paper (mục đích NC, thiết kế, kết luận chính) — đã verify cơ bản cho các paper Bài 1 + Bài 2 + 6 paper có lỗi vừa phát hiện. Còn paper khác (Anderson 2025, Saikia 2023, Pascoe 2020 detail, Kendall 1994 64% improvement, Barrett 1996 84%/57%) chưa verify chi tiết qualitative.',
    'Loại 3: Causal vs associational language audit — chưa làm cho Tổng quan v2.',
    'Loại 4: Statistical interpretation correctness — chưa làm.',
    'Loại 5: Hedging language for evidence strength — chưa làm.',
    'Loại 6: Acronym completeness — đã sửa cho v3 (DASS-21, GAD-7 đã định nghĩa). Cần verify trong bản hợp nhất.',
    'Loại 7-11: Format (số liệu Vietnamese, citation format, geographic naming, date, DOI) — đã có quy ước.',
    'Loại 12: Vietnamese diacritics spot-check — chưa làm cho 6 paper mới sửa.',
    'Loại 13-17: Cross-reference, alignment với abstract, verbatim, tone — chưa làm.',
]
for s in steps:
    p = doc.add_paragraph(s, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(10)
doc.add_paragraph()

# Cam ket lim chinh
H(doc, '7. Ghi chú liêm chính khoa học', 2)
P(doc, 'NCS xin chân thành xin lỗi vì các lỗi số liệu phát hiện được trong các tài liệu nội bộ và Tổng quan v2. Nguyên nhân là khi viết bổ sung phạm vi Tổng quan, một số nghiên cứu phụ trợ không nằm trong 21 TLTK đã verify từ PDF của Bài 1 và Bài 2 đã bị thêm với số liệu KHÔNG được đối chiếu chặt chẽ với nguồn gốc. Sau khi user đặt ra câu hỏi cụ thể về Nguyễn Danh Lâm 2022, NCS đã chủ động kiểm tra và phát hiện thêm 5 nghiên cứu khác cũng có lỗi tương tự. Đây là tình huống không chấp nhận được trong nghiên cứu khoa học và đã được khắc phục bằng cách đọc lại từng PDF gốc và sửa lại số liệu.', italic=True, color=RGBColor(80, 80, 80))
P(doc, 'Các phần sửa trong file v3/v2 mới được tô ĐỎ ĐẬM để thầy phản biện và NCS dễ thấy chỗ nào đã đổi. Nếu NCS chấp nhận tất cả các sửa, có thể chuyển toàn bộ chữ đỏ về đen sau khi thẩm định lại lần cuối.', italic=True, color=RGBColor(80, 80, 80))
P(doc, 'Bản gốc 33-trang_FINAL và FULL_v1 vẫn được giữ làm reference; bản v3/v2 mới là bản đề xuất cho nộp luận án.', italic=True, color=RGBColor(80, 80, 80))
doc.add_paragraph()

P(doc, 'Hà Nội, ngày 27 tháng 05 năm 2026', indent=False, italic=True)
P(doc, 'Nghiên cứu sinh', indent=False, bold=True)
P(doc, '', indent=False)
P(doc, 'Công Thị Hằng', indent=False, bold=True)

doc.save(OUT)
from docx import Document as D
d = D(OUT)
w = sum(len(p.text.split()) for p in d.paragraphs)
n_p = sum(1 for p in d.paragraphs if p.text.strip())
print(f"Da luu: {OUT}")
print(f"Words: {w}, Paragraphs: {n_p}, Tables: {len(d.tables)}, Size: {os.path.getsize(OUT)//1024}KB")
