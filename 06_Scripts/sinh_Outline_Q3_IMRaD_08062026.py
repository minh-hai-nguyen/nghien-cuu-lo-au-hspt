# -*- coding: utf-8 -*-
"""Outline IMRaD bai Q3 — Multi-group SEM theo gioi.
Song ngu Anh-Viet, ~3000 tu, dung lai material tu Q2 v1 + 6 luan chung
trong TraLoiThayNMD."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'Outline_Q3_IMRaD_SongNgu_v1_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def TITLE(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def EN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def VN(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(8)
    p.paragraph_format.first_line_indent = Cm(0.5); p.paragraph_format.left_indent = Cm(0.3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def PB(en, vn):
    EN(en); VN(vn)

def BB(en, vn, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('• ' + en); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8 + level * 0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('▸ ' + vn); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True
    r.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

def NOTE(en, vn):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('[NOTE TO NCS + co-authors] ' + en)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run('[GHI CHÚ cho NCS + đồng tác giả] ' + vn)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True; r.italic = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)


# =====================================================
TITLE('OUTLINE IMRaD — BÀI Q3')
TITLE('Gender-specific pathways to anxiety disorders in early '
      'adolescence: A multi-group SEM in Vietnamese lower '
      'secondary students')
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Đường dẫn đặc thù theo giới đến các rối loạn lo âu ở đầu '
              'vị thành niên: Mô hình phương trình cấu trúc đa nhóm '
              'trên học sinh trung học cơ sở Việt Nam')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Outline song ngữ Anh-Việt v1 — soạn 08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

NOTE('This is an IMRaD outline (structure + key claims, not full prose). '
     'Authors: Cong TH (1st), [Nguyen MD], Hai-Nguyen Minh. Target journal: '
     'Frontiers in Psychiatry (Adolescent and Young Adult Psychiatry '
     'section) OR BMC Public Health. Word target full draft: 7,000-9,000.',
     'Đây là outline IMRaD (cấu trúc + luận điểm chính, không phải bản '
     'thảo đầy đủ). Tác giả: Công TH (chủ trì), [Nguyễn MD], Hải-Nguyễn '
     'Minh. Tạp chí đích: Frontiers in Psychiatry (section Adolescent '
     'and Young Adult Psychiatry) HOẶC BMC Public Health. Mục tiêu '
     'word count cho draft đầy đủ: 7.000-9.000 từ.')


# =====================================================
H1('ABSTRACT (300 words target)')

H2('Background')
PB('Adolescent anxiety disorders are more common in girls than boys '
   'globally, but the magnitude of this gender gap varies by '
   'developmental stage and cultural context. The largest cross-'
   'national study to date (Jefferies & Ungar 2020, N = 6,825 in 7 '
   'countries) found no overall gender difference in social anxiety '
   'among young adults aged 16-29, yet our preliminary data in '
   'Vietnamese lower secondary students show female-higher scores '
   'on Social Anxiety Disorder (SAD). This raises a key research '
   'question: are the entire risk-protective SEM pathways invariant '
   'across gender in early adolescence?',
   'Các rối loạn lo âu ở vị thành niên xảy ra phổ biến hơn ở nữ so '
   'với nam trên toàn cầu, nhưng độ lớn của khác biệt này thay đổi '
   'theo giai đoạn phát triển và bối cảnh văn hóa. Nghiên cứu xuyên '
   'quốc gia lớn nhất đến nay (Jefferies & Ungar 2020, N = 6.825 ở '
   '7 quốc gia) không phát hiện khác biệt giới về lo âu xã hội ở '
   'thanh niên 16-29 tuổi, nhưng dữ liệu sơ bộ của chúng tôi trên '
   'học sinh trung học cơ sở Việt Nam cho thấy điểm SAD của nữ cao '
   'hơn nam. Câu hỏi nghiên cứu cốt lõi: liệu toàn bộ các đường dẫn '
   'của mô hình SEM nguy cơ-bảo vệ có bất biến theo giới ở đầu vị '
   'thành niên không?')

H2('Methods')
PB('Cross-sectional survey of 1,352 Vietnamese lower secondary '
   'students (614 boys, 738 girls; ages 11-14) from two purposively '
   'selected schools in Hanoi (Nhat Tan urban, Tay Mo suburban). '
   'Eight validated scales measured three risk factors '
   '(bullying, academic stress, smartphone addiction), four '
   'protective factors (school membership, parental support, peer '
   'support, self-esteem), and three anxiety subtypes (GAD, SAD, '
   'SocAD). Multi-group SEM tested configural, metric, and scalar '
   'measurement invariance, then structural invariance of all 21 '
   'pathways, using ΔCFI ≤ 0.01 (Cheung & Rensvold 2002) and '
   'ΔRMSEA ≤ 0.015.',
   'Khảo sát cắt ngang 1.352 học sinh THCS Việt Nam (614 nam, 738 '
   'nữ; 11-14 tuổi) từ hai trường được chọn có chủ đích tại Hà Nội '
   '(Nhật Tân — đô thị, Tây Mỗ — ngoại ô). Tám thang đo đã được '
   'chuẩn hóa đo ba yếu tố nguy cơ (bắt nạt, áp lực học tập, '
   'nghiện điện thoại), bốn yếu tố bảo vệ (gắn bó trường, hỗ trợ '
   'cha mẹ, hỗ trợ bạn bè, tự trọng), và ba phân loại lo âu (GAD, '
   'SAD, SocAD). SEM đa nhóm kiểm chứng bất biến đo lường ở cấp '
   'cấu hình, metric, scalar, sau đó kiểm chứng bất biến cấu trúc '
   'của toàn bộ 21 đường dẫn, sử dụng ΔCFI ≤ 0,01 (Cheung & '
   'Rensvold 2002) và ΔRMSEA ≤ 0,015.')

H2('Results (planned)')
PB('We expect to confirm measurement invariance at all three levels '
   '(supporting valid gender comparison). For structural invariance, '
   'we anticipate identifying a subset of non-invariant pathways — '
   'particularly bullying → SAD (β stronger in girls per preliminary '
   'data, p < .001) and possibly self-esteem → GAD. Effect sizes '
   'and 95% CIs will be reported for all 21 pathways stratified by '
   'gender.',
   'Chúng tôi kỳ vọng khẳng định bất biến đo lường ở cả ba cấp (hỗ '
   'trợ so sánh giới hợp lệ). Với bất biến cấu trúc, dự đoán nhận '
   'diện được một tập con các đường dẫn không bất biến — đặc biệt '
   'bắt nạt → SAD (β mạnh hơn ở nữ theo dữ liệu sơ bộ, p < 0,001) '
   'và có thể tự trọng → GAD. Hệ số effect size và CI 95% sẽ được '
   'báo cáo cho cả 21 đường dẫn phân tầng theo giới.')

H2('Conclusions (planned)')
PB('The findings will inform gender-tailored prevention programs '
   'and reconcile the discrepancy between our sample and Jefferies '
   '& Ungar (2020) via a developmental-cultural framework drawing '
   'on Hankin (2007), McLean & Anderson (2009), Rose (2002), '
   'Stankov (2010), and Small & Blanc (2021).',
   'Các phát hiện sẽ cung cấp thông tin cho chương trình phòng '
   'ngừa đặc thù giới và lý giải sự không nhất quán giữa mẫu '
   'nghiên cứu của chúng tôi với Jefferies & Ungar (2020) thông '
   'qua khung phát triển-văn hóa dựa trên Hankin (2007), McLean & '
   'Anderson (2009), Rose (2002), Stankov (2010), Small & Blanc '
   '(2021).')


# =====================================================
H1('1. INTRODUCTION (target 1,000-1,200 words full draft)')

H2('1.1 — The gender paradox in adolescent anxiety')
PB('Open with the global epidemiological finding: girls show higher '
   'rates of anxiety disorders than boys in most international '
   'studies, with the gap emerging around puberty (McLean & '
   'Anderson 2009; Hankin et al. 2007). Cite Kessler et al. 2005 '
   'on age-of-onset (half of all anxiety begins by age 14). State '
   'the puzzle: this finding does NOT replicate in some recent '
   'large-N cross-national studies (Jefferies & Ungar 2020 — N = '
   '6,825 in 7 countries, no overall gender effect).',
   'Mở đầu bằng phát hiện dịch tễ toàn cầu: nữ có tỷ lệ mắc rối '
   'loạn lo âu cao hơn nam trong hầu hết nghiên cứu quốc tế, '
   'khoảng cách xuất hiện quanh tuổi dậy thì (McLean & Anderson '
   '2009; Hankin và cs. 2007). Trích Kessler và cs. 2005 về độ '
   'tuổi khởi phát (một nửa số ca lo âu khởi phát trước 14 tuổi). '
   'Nêu nghịch lý: phát hiện này KHÔNG được lặp lại trong một số '
   'nghiên cứu xuyên quốc gia gần đây có cỡ mẫu lớn (Jefferies & '
   'Ungar 2020 — N = 6.825 ở 7 quốc gia, không có hiệu ứng giới '
   'tổng thể).')

H2('1.2 — Three explanations for the discrepancy')
PB('Briefly preview three theoretical mechanisms that may explain '
   'why gender differences appear in early adolescence (11-14) but '
   'disappear in young adulthood (16-29): (1) developmental peak '
   'vulnerability — gender role intensification + co-rumination '
   'peak in middle school years; (2) cultural moderation — '
   'Confucian academic culture imposes stricter gender expectations '
   'on Vietnamese girls than Western contexts (Stankov 2010); (3) '
   'methodological factors — measurement context, sampling bias, '
   'and age-pooling effects.',
   'Giới thiệu ngắn ba cơ chế lý thuyết có thể giải thích vì sao '
   'khác biệt giới xuất hiện ở đầu vị thành niên (11-14) nhưng '
   'biến mất ở thanh niên (16-29): (1) đỉnh điểm tổn thương phát '
   'triển — gender intensification + co-rumination đạt đỉnh ở tuổi '
   'học sinh THCS; (2) điều tiết văn hóa — văn hóa học thuật Nho '
   'giáo áp đặt kỳ vọng giới nghiêm ngặt hơn cho nữ Việt Nam so '
   'với bối cảnh phương Tây (Stankov 2010); (3) yếu tố phương '
   'pháp luận — bối cảnh đo lường, sai lệch chọn mẫu, và hiệu ứng '
   'pool tuổi.')

H2('1.3 — Why multi-group SEM is the right tool')
PB('Conventional t-test or ANOVA on total anxiety scores can '
   'detect mean differences but cannot test whether the entire '
   'risk-protective MECHANISM differs by gender. Multi-group SEM '
   'tests measurement invariance (do scales measure the same '
   'construct in both groups?) and structural invariance (do '
   'pathway coefficients differ?) — providing a fine-grained map '
   'of where gender matters and where it does not.',
   'Kiểm định t-test hay ANOVA thông thường trên điểm lo âu tổng '
   'có thể phát hiện khác biệt trung bình nhưng không kiểm chứng '
   'được liệu toàn bộ CƠ CHẾ nguy cơ-bảo vệ có khác biệt theo giới '
   'không. SEM đa nhóm kiểm chứng bất biến đo lường (các thang đo '
   'có đo cùng cấu trúc ở cả hai nhóm không?) và bất biến cấu trúc '
   '(các hệ số đường dẫn có khác nhau không?) — cung cấp bản đồ '
   'chi tiết về nơi giới có vai trò và nơi không.')

H2('1.4 — Present study and hypotheses')
PB('State three pre-registered hypotheses. H1: Measurement '
   'invariance will be supported at all three levels (configural, '
   'metric, scalar) for the 8-scale measurement model. H2: At '
   'least one structural pathway will be non-invariant, specifically '
   'the bullying → SAD pathway will be stronger in girls than boys '
   '(prior cohort evidence). H3: The integrated higher-order model '
   '(YTNC + YTBV → total RLLA, R² = 0.598 in pooled sample) will '
   'show different total effect sizes by gender.',
   'Nêu ba giả thuyết đã đăng ký trước. H1: Bất biến đo lường sẽ '
   'được hỗ trợ ở cả ba cấp (cấu hình, metric, scalar) cho mô '
   'hình đo lường 8 thang. H2: Ít nhất một đường dẫn cấu trúc sẽ '
   'không bất biến, cụ thể đường dẫn bắt nạt → SAD sẽ mạnh hơn ở '
   'nữ so với nam (bằng chứng cohort trước đây). H3: Mô hình bậc '
   'cao tích hợp (YTNC + YTBV → tổng RLLA, R² = 0,598 ở mẫu '
   'pooled) sẽ thể hiện effect size tổng khác nhau theo giới.')

NOTE('Hypotheses H1-H3 should be confirmed by NCS before pre-'
     'registration. Pre-registration platform: OSF or AsPredicted. '
     'Pre-registration is a prerequisite for top-tier Q1 anxiety '
     'journals.',
     'Giả thuyết H1-H3 cần được NCS xác nhận trước khi đăng ký '
     'trước. Nền tảng đăng ký trước: OSF hoặc AsPredicted. Đăng '
     'ký trước là yêu cầu của các tạp chí Q1 hàng đầu về lo âu.')


# =====================================================
H1('2. METHODS (target 1,000-1,200 words full draft)')

H2('2.1 — Participants')
PB('Reuse Q2 v1 participant description verbatim: 1,352 lower '
   'secondary students (boys = 614, 45.4%; girls = 738, 54.6%) '
   'recruited from Nhat Tan (urban) and Tay Mo (suburban) schools '
   'in Hanoi. Ages 11-14 (grades 6-9 distribution: 368/316/340/'
   '328). Written parental consent + student assent obtained.',
   'Sử dụng lại phần Participants của Q2 v1 nguyên văn: 1.352 học '
   'sinh THCS (nam = 614, 45,4%; nữ = 738, 54,6%) tuyển từ trường '
   'Nhật Tân (đô thị) và Tây Mỗ (ngoại ô) tại Hà Nội. Tuổi 11-14 '
   '(phân bố khối 6-9: 368/316/340/328). Có chấp thuận viết của '
   'cha mẹ + đồng thuận viết của học sinh.')

H2('2.2 — Measures')
PB('Reuse Q2 v1 measures section verbatim — eight validated scales '
   '(RCADS for GAD/SAD/SocAD; OBVQ for bullying; ESSA for academic '
   'stress; SAS-SV for smartphone addiction; PSSM for school '
   'membership; MSPSS for parental + peer support; RSES for self-'
   'esteem). All scores rescaled to 0-100 metric for cross-scale '
   'comparability while preserving original scoring conventions.',
   'Sử dụng lại phần Measures của Q2 v1 nguyên văn — tám thang đo '
   'đã chuẩn hóa (RCADS cho GAD/SAD/SocAD; OBVQ cho bắt nạt; ESSA '
   'cho áp lực học tập; SAS-SV cho nghiện điện thoại; PSSM cho '
   'gắn bó trường; MSPSS cho hỗ trợ cha mẹ + bạn bè; RSES cho '
   'tự trọng). Tất cả điểm được chuyển sang thang 0-100 để so '
   'sánh xuyên thang trong khi vẫn giữ quy ước tính điểm gốc.')

H2('2.3 — Analytic strategy (Q3-specific, NEW)')

BB('Step 1: Descriptive statistics + bivariate gender comparisons. '
   'Independent t-tests for continuous scales (with Welch correction '
   'for unequal variances); chi-square for SAD threshold. Effect '
   'sizes (Cohen\'s d) and 95% CIs reported.',
   'Bước 1: Thống kê mô tả + so sánh giới bivariate. Kiểm định '
   't-test độc lập cho các thang liên tục (với hiệu chỉnh Welch '
   'nếu phương sai không đồng nhất); chi-square cho ngưỡng SAD. '
   'Effect size (Cohen\'s d) và CI 95% được báo cáo.', 0)

BB('Step 2: Measurement invariance testing via multi-group CFA. '
   'Sequence: configural → metric → scalar. Criterion: ΔCFI ≤ '
   '0.01 + ΔRMSEA ≤ 0.015 (Cheung & Rensvold 2002; Chen 2007). '
   'Partial invariance acceptable for ≤ 20% of indicators per '
   'scale.',
   'Bước 2: Kiểm chứng bất biến đo lường qua CFA đa nhóm. Trình '
   'tự: cấu hình → metric → scalar. Tiêu chí: ΔCFI ≤ 0,01 + '
   'ΔRMSEA ≤ 0,015 (Cheung & Rensvold 2002; Chen 2007). Bất '
   'biến từng phần chấp nhận được nếu ≤ 20% chỉ báo trong mỗi '
   'thang.', 0)

BB('Step 3: Structural invariance testing. Compare unconstrained '
   '(all 21 pathways free) vs constrained (all 21 pathways equal '
   'across groups) models. Identify which pathways need to be '
   'released for partial invariance.',
   'Bước 3: Kiểm chứng bất biến cấu trúc. So sánh mô hình '
   'không ràng buộc (cả 21 đường dẫn tự do) vs ràng buộc (cả 21 '
   'đường dẫn bằng nhau qua các nhóm). Xác định đường dẫn nào '
   'cần được nới để đạt bất biến từng phần.', 0)

BB('Step 4: Higher-order model invariance. Test invariance of '
   'YTNC + YTBV second-order pathways to total RLLA (Q2 v1 '
   'integrated model). Report β by gender + 95% CIs.',
   'Bước 4: Bất biến của mô hình bậc cao. Kiểm chứng bất biến '
   'các đường dẫn bậc hai YTNC + YTBV → tổng RLLA (mô hình '
   'tích hợp của Q2 v1). Báo cáo β theo giới + CI 95%.', 0)

BB('Step 5: Sensitivity analyses. Multi-group SEM repeated '
   'after (a) excluding urban-only / suburban-only sub-samples, '
   '(b) age stratification (younger 11-12 vs older 13-14), (c) '
   'alternative invariance threshold (Chen 2007 stricter).',
   'Bước 5: Phân tích độ nhạy. Lặp lại SEM đa nhóm sau khi (a) '
   'loại trừ con mẫu chỉ-đô-thị / chỉ-ngoại-ô, (b) phân tầng '
   'theo tuổi (nhỏ hơn 11-12 vs lớn hơn 13-14), (c) ngưỡng bất '
   'biến thay thế (Chen 2007 nghiêm hơn).', 0)

H2('2.4 — Ethics')
PB('Reuse Q2 v1 ethics section verbatim. [TBD — Q3-6 BLOCKING: '
   'IRB number and date pending official letter from HNUE Ethics '
   'Committee — NCS Cong Thi Hang to confirm.]',
   'Sử dụng lại phần Ethics của Q2 v1 nguyên văn. [CHỜ XÁC NHẬN '
   '— Q3-6 BLOCKING: số quyết định và ngày của HĐ đạo đức HNUE '
   'chờ thư chính thức — NCS Công Thị Hằng xác nhận.]')


# =====================================================
H1('3. RESULTS (target 1,500-1,800 words full draft)')

H2('3.1 — Descriptive statistics and bivariate gender comparisons')
PB('Report means, SDs, and Cronbach alphas for each of 8 scales, '
   'stratified by gender. Highlight which scales show significant '
   'mean differences (expected: SAD, SocAD, self-esteem). Include '
   'Table 1 (gender x scale descriptives).',
   'Báo cáo trung bình, SD, và alpha Cronbach cho mỗi trong 8 '
   'thang, phân tầng theo giới. Nhấn mạnh thang nào có khác biệt '
   'trung bình có ý nghĩa (dự kiến: SAD, SocAD, tự trọng). Bao '
   'gồm Bảng 1 (descriptives giới × thang).')

H2('3.2 — Measurement invariance')
PB('Report sequential model fit indices: CFI, TLI, RMSEA, SRMR '
   'for configural, metric, scalar models. Report ΔCFI, ΔRMSEA. '
   'Discuss any partial invariance findings (which items released).',
   'Báo cáo chỉ số phù hợp mô hình tuần tự: CFI, TLI, RMSEA, '
   'SRMR cho mô hình cấu hình, metric, scalar. Báo cáo ΔCFI, '
   'ΔRMSEA. Bàn luận bất kỳ phát hiện bất biến từng phần (mục '
   'nào được nới).')

H2('3.3 — Structural invariance')
PB('Compare unconstrained vs constrained model. Identify which '
   'of 21 pathways are non-invariant. Present Table 2 (β by '
   'gender for each pathway + 95% CIs + z-test for invariance).',
   'So sánh mô hình không ràng buộc vs ràng buộc. Xác định '
   'đường dẫn nào trong 21 đường không bất biến. Trình bày Bảng '
   '2 (β theo giới cho mỗi đường dẫn + CI 95% + kiểm định z '
   'bất biến).')

H2('3.4 — Pathway-specific gender effects')
PB('Focus on substantively interesting non-invariant pathways. '
   'Expected key finding (TO BE CONFIRMED by analysis — NOT a '
   'fabricated estimate): bullying → SAD β may be stronger in '
   'girls than boys, but exact β values await multi-group SEM '
   'computation. Discuss effect size magnitude in narrative form.',
   'Tập trung vào các đường dẫn không bất biến có ý nghĩa thực '
   'chất. Phát hiện chính dự kiến (CHỜ XÁC NHẬN từ phân tích — '
   'KHÔNG phải ước tính bịa): đường dẫn bắt nạt → SAD có thể '
   'mạnh hơn ở nữ so với nam, nhưng giá trị β chính xác chờ '
   'kết quả tính toán SEM đa nhóm. Bàn luận độ lớn effect size '
   'qua mô tả ngôn từ.')

H2('3.5 — Higher-order integrated model')
PB('Report YTNC and YTBV second-order pathway β by gender. R² '
   'comparison.',
   'Báo cáo β đường dẫn bậc hai YTNC và YTBV theo giới. So sánh '
   'R².')


# =====================================================
H1('4. DISCUSSION (target 1,500-1,800 words full draft)')

H2('4.1 — Summary of findings')
PB('1 paragraph summary of measurement + structural invariance + '
   'key non-invariant pathways.',
   '1 đoạn tóm tắt bất biến đo lường + cấu trúc + đường dẫn '
   'không bất biến chính.')

H2('4.2 — Reconciliation with Jefferies & Ungar (2020)')
PB('THE KEY DISCUSSION SECTION. Reuse the 6 luan chung from '
   'TraLoiThayNMD doc verbatim (8 paragraphs total, ~600 words). '
   'Sequence: Lc1 (developmental peak 11-14, Hankin 2007, McLean '
   '& Anderson 2009) → Lc2 (co-rumination, Rose 2002) → Lc3 '
   '(Confucian gender script, Stankov 2010, Small & Blanc 2021) '
   '→ Lc4 (school context vs market panel) → Lc5 (methodological: '
   'instrument, sample bias, sub-sample size) → Lc6 (age main '
   'effect dominates Jefferies).',
   'SECTION DISCUSSION CHỦ CHỐT. Sử dụng lại 6 luận chứng từ '
   'doc TraLoiThayNMD nguyên văn (tổng 8 đoạn, ~600 từ). Trình '
   'tự: Lc1 (đỉnh phát triển 11-14, Hankin 2007, McLean & '
   'Anderson 2009) → Lc2 (co-rumination, Rose 2002) → Lc3 '
   '(kịch bản giới Nho giáo, Stankov 2010, Small & Blanc 2021) '
   '→ Lc4 (bối cảnh học đường vs market panel) → Lc5 (phương '
   'pháp luận: công cụ, sai lệch mẫu, cỡ con mẫu) → Lc6 (hiệu '
   'ứng tuổi áp đảo trong Jefferies).')

H2('4.3 — Theoretical implications')
PB('Three theoretical contributions: (1) gender intensification '
   'theory (Hill & Lynch 1983) applies to non-Western Confucian '
   'context — extends generalizability; (2) co-rumination cuts '
   'across cultures but interacts with Confucian disclosure norms '
   '(Small & Blanc 2021); (3) tam giao framework (Buddhism + '
   'Confucianism + Taoism) provides an under-theorized model of '
   'emotional restraint relevant to gendered SAD expression.',
   'Ba đóng góp lý thuyết: (1) lý thuyết gender intensification '
   '(Hill & Lynch 1983) áp dụng cho bối cảnh Nho giáo phi phương '
   'Tây — mở rộng tính tổng quát; (2) co-rumination xuyên văn hóa '
   'nhưng tương tác với chuẩn mực bộc lộ Nho giáo (Small & Blanc '
   '2021); (3) khung tam giáo (Phật + Nho + Đạo) cung cấp mô '
   'hình chừng mực cảm xúc chưa được lý thuyết hóa đầy đủ, liên '
   'quan đến biểu hiện SAD theo giới.')

H2('4.4 — Practical implications')
PB('Gender-tailored prevention: girls in early adolescence may '
   'benefit from co-rumination disruption + assertiveness training '
   '+ self-compassion. Boys may benefit from emotion identification '
   'and labeling. School-based programs (e.g., FRIENDS, OurFutures) '
   'should be tested in gender-stratified RCTs in Vietnamese '
   'schools.',
   'Phòng ngừa đặc thù theo giới: nữ ở đầu vị thành niên có thể '
   'hưởng lợi từ việc làm gián đoạn co-rumination + huấn luyện '
   'quyết đoán + tự thương xót. Nam có thể hưởng lợi từ nhận '
   'diện + đặt tên cảm xúc. Chương trình tại trường (vd FRIENDS, '
   'OurFutures) nên được kiểm chứng trong RCT phân tầng giới ở '
   'trường học Việt Nam.')

H2('4.5 — Limitations')
PB('(a) Cross-sectional design precludes causal inference; (b) '
   'single-site Hanoi sample limits generalization to rural / '
   'minority-ethnic Vietnamese youth; (c) self-report instruments '
   'subject to social desirability + gender-biased disclosure '
   'norms (which is itself a finding of interest); (d) single-'
   'time-point measurement does not capture developmental '
   'trajectories.',
   '(a) Thiết kế cắt ngang không cho phép suy luận nhân quả; '
   '(b) mẫu một địa điểm Hà Nội hạn chế khái quát hóa cho thanh '
   'thiếu niên nông thôn / dân tộc thiểu số Việt Nam; (c) công '
   'cụ tự báo cáo bị ảnh hưởng bởi mong muốn xã hội + chuẩn mực '
   'bộc lộ thiên lệch giới (bản thân điều này cũng là phát hiện '
   'đáng quan tâm); (d) đo lường một thời điểm không bắt được '
   'quỹ đạo phát triển.')

H2('4.6 — Future directions')
PB('Three priorities: (1) longitudinal cohort following the same '
   'sample through transition to upper secondary (16-18) to test '
   'whether gender gap closes — direct test of developmental peak '
   'hypothesis; (2) replicate in rural + minority-ethnic Vietnamese '
   'cohorts; (3) qualitative companion study probing gendered '
   'disclosure norms (planned, not in current dataset).',
   'Ba ưu tiên: (1) cohort theo dõi mẫu hiện tại qua chuyển tiếp '
   'lên THPT (16-18) để kiểm chứng xem khoảng cách giới có khép '
   'lại không — kiểm định trực tiếp giả thuyết đỉnh phát triển; '
   '(2) lặp lại ở cohort nông thôn + dân tộc thiểu số Việt Nam; '
   '(3) nghiên cứu định tính đồng hành thăm dò chuẩn mực bộc lộ '
   'theo giới (đã lên kế hoạch, không có trong dữ liệu hiện tại).')


# =====================================================
H1('5. REFERENCES (target 35-45 entries full draft)')

NOTE('Below are 12 KEY references already verified via PubMed (PMID + '
     'DOI checked 07-08/06/2026). For full draft, expand to 35-45 by '
     'adding instrument validation papers (RCADS Chorpita 2000, OBVQ '
     'Olweus 1996, etc.) and additional invariance methodology refs.',
     'Dưới đây là 12 tham khảo CHỦ CHỐT đã verify qua PubMed (PMID + '
     'DOI kiểm tra 07-08/06/2026). Cho draft đầy đủ, mở rộng lên 35-45 '
     'bằng cách thêm các bài chuẩn hóa thang đo (RCADS Chorpita 2000, '
     'OBVQ Olweus 1996, v.v.) và các tham khảo phương pháp luận bất '
     'biến bổ sung.')

refs = [
    '[1] Jefferies P, Ungar M. (2020). Social anxiety in young '
    'people: A prevalence study in seven countries. PLOS ONE '
    '15(9):e0239133. DOI: 10.1371/journal.pone.0239133.',
    '[2] Hankin BL, Mermelstein R, Roesch L. (2007). Sex '
    'differences in adolescent depression: Stress exposure and '
    'reactivity models. Child Development 78(1):279-295. '
    'PMID: 17328705.',
    '[3] McLean CP, Anderson ER. (2009). Brave men and timid '
    'women? A review of the gender differences in fear and '
    'anxiety. Clinical Psychology Review 29(6):496-505. '
    'PMID: 19541399.',
    '[4] Hill JP, Lynch ME. (1983). The intensification of '
    'gender-related role expectations during early adolescence. '
    'In: Brooks-Gunn J, Petersen AC (eds.), Girls at puberty. '
    'New York: Plenum Press, pp. 201-228.',
    '[5] Kessler RC, Berglund P, Demler O, Jin R, Merikangas KR, '
    'Walters EE. (2005). Lifetime prevalence and age-of-onset '
    'distributions of DSM-IV disorders in the NCS-R. Archives of '
    'General Psychiatry 62(6):593-602. PMID: 15939837. '
    'DOI: 10.1001/archpsyc.62.6.593.',
    '[6] Rose AJ. (2002). Co-rumination in the friendships of '
    'girls and boys. Child Development 73(6):1830-1843. '
    'PMID: 12487497. DOI: 10.1111/1467-8624.00509.',
    '[7] Stankov L. (2010). Unforgiving Confucian culture: A '
    'breeding ground for high academic achievement, test anxiety '
    'and self-doubt? Learning and Individual Differences '
    '20(6):555-563. DOI: 10.1016/j.lindif.2010.05.003.',
    '[8] Small S, Blanc J. (2021). Mental Health During COVID-19: '
    'Tam Giao and Vietnam\'s Response. Frontiers in Psychiatry '
    '11:589618. PMID: 33536961. DOI: 10.3389/fpsyt.2020.589618.',
    '[9] Cheung GW, Rensvold RB. (2002). Evaluating goodness-of-'
    'fit indexes for testing measurement invariance. Structural '
    'Equation Modeling 9(2):233-255.',
    '[10] Chen FF. (2007). Sensitivity of goodness of fit indexes '
    'to lack of measurement invariance. Structural Equation '
    'Modeling 14(3):464-504.',
    '[11] Hu LT, Bentler PM. (1999). Cutoff criteria for fit '
    'indexes in covariance structure analysis. Structural '
    'Equation Modeling 6(1):1-55.',
    '[12] Compas BE, Jaser SS, Bettis AH, et al. (2017). Coping, '
    'emotion regulation, and psychopathology in childhood and '
    'adolescence. Psychological Bulletin 143(9):939-991. '
    'PMID: 28616996. DOI: 10.1037/bul0000110.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Footer
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
r = p.add_run('Soạn 08/06/2026 — Outline Q3 v1, chờ NCS + thầy NMĐ confirm '
              'trước khi viết full draft')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')

from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
for t in d2.tables:
    for row in t.rows:
        for cell in row.cells:
            chunks.append(cell.text)
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
