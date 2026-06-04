# -*- coding: utf-8 -*-
"""Kiem tra ky tung dong: TOC table + body content.
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

print("="*80)
print("KIEM TRA CHI TIET TUNG DONG")
print("="*80)

# ============================================================
# 1. FULL TOC table content
# ============================================================
print("\n[1] TOAN BO BANG TOC TABLE")
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and 'MỞ ĐẦU' in tb.rows[0].cells[0].text:
        toc_table = tb; break
if toc_table:
    print(f"  Total: {len(toc_table.rows)} rows")
    print()
    for i, row in enumerate(toc_table.rows):
        name = row.cells[0].text.strip()
        page = row.cells[1].text.strip() if len(row.cells) > 1 else ''
        print(f"  {i:>3}: {name[:75]:<75} | {page}")

# ============================================================
# 2. Check all Heading-styled paras in doc - any missing from TOC?
# ============================================================
print()
print("="*80)
print("[2] CAC HEADING TRONG DOC - MATCH TOC")
print("="*80)

# Collect heading paragraphs in doc (excluding tables)
def in_table(p):
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag in (qn('w:tbl'), qn('w:tc')):
            return True
        parent = parent.getparent()
    return False

heading_paras = []
for i, p in enumerate(d.paragraphs):
    if in_table(p): continue
    style = p.style.name if p.style else ''
    if 'Heading' in style or 'Tiêu đề' in style:
        text = p.text.strip()
        if not text or len(text) > 250: continue
        level_m = re.search(r'\d+', style)
        level = int(level_m.group()) if level_m else 0
        heading_paras.append((i, level, text, style))

print(f"\n  Total heading paragraphs in doc: {len(heading_paras)}")
print()
# Show H1-H3 only (matches TOC)
h1_3 = [h for h in heading_paras if h[1] <= 3]
print(f"  H1-H3 (TOC level): {len(h1_3)}")

# ============================================================
# 3. Check TOC <-> heading consistency
# ============================================================
print()
print("="*80)
print("[3] KIEM TRA NHAT QUAN TOC <-> DOC HEADINGS")
print("="*80)

def normalize(s):
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'^[\d\.]+\s*', '', s)
    return s.strip(' .')

toc_names_norm = set()
for row in toc_table.rows:
    name = row.cells[0].text.strip()
    norm = normalize(name)
    toc_names_norm.add(norm)

heading_norms = set()
for i, lv, text, style in heading_paras:
    if lv > 3: continue
    norm = normalize(text)
    heading_norms.add(norm)

# Find ones in heading but not in TOC
in_heading_not_toc = []
for i, lv, text, style in heading_paras:
    if lv > 3: continue
    norm = normalize(text)
    if norm in toc_names_norm: continue
    # Skip obvious skips
    if text.upper().strip() in {'MỤC LỤC', 'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
                                  'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
                                  'LỜI CAM ĐOAN', 'LỜI CẢM ƠN'}:
        continue
    # Skip CHƯƠNG X (combined into chapter title in TOC)
    if re.match(r'^CHƯƠNG\s+\d+$', text): continue
    in_heading_not_toc.append((i, lv, text))

print(f"\n  Headings in doc nhung CHUA co trong TOC: {len(in_heading_not_toc)}")
for i, lv, text in in_heading_not_toc[:15]:
    print(f"    para {i} L{lv}: {text[:80]}")

# Find ones in TOC but no matching heading
in_toc_not_heading = []
for row in toc_table.rows:
    name = row.cells[0].text.strip()
    norm = normalize(name)
    if norm in heading_norms: continue
    # Skip MỞ ĐẦU (which IS a heading)
    if name.upper().strip() in {'MỞ ĐẦU', 'KẾT LUẬN VÀ KIẾN NGHỊ', 'TÀI LIỆU THAM KHẢO', 'PHỤ LỤC'}:
        continue
    # Skip CHƯƠNG X title - we combine these in TOC
    if 'CHƯƠNG' in name.upper() and ':' in name:
        # Extract title after colon, check if exists
        title = name.split(':', 1)[1].strip()
        if normalize(title) in heading_norms: continue
    in_toc_not_heading.append(name)

print(f"\n  TOC entries KHONG match heading nao: {len(in_toc_not_heading)}")
for name in in_toc_not_heading[:15]:
    print(f"    {name[:80]}")

# ============================================================
# 4. Check duplicate paragraphs (might indicate doubled content)
# ============================================================
print()
print("="*80)
print("[4] KIEM TRA DOAN TRUNG LAP (co the la noi dung bi double)")
print("="*80)

para_text_count = {}
for i, p in enumerate(d.paragraphs):
    text = p.text.strip()
    if len(text) < 30: continue  # skip short ones
    if text in para_text_count:
        para_text_count[text].append(i)
    else:
        para_text_count[text] = [i]

dupes = {k: v for k, v in para_text_count.items() if len(v) > 1}
print(f"\n  Doan trung lap (>= 30 ky tu): {len(dupes)}")
for text, paras in list(dupes.items())[:5]:
    print(f"    Para {paras}: {text[:80]}")

# ============================================================
# 5. Statistics
# ============================================================
print()
print("="*80)
print("[5] STATISTICS")
print("="*80)
print(f"  Paragraphs: {len(d.paragraphs)}")
print(f"  Tables: {len(d.tables)}")
print(f"  Total chars: {sum(len(p.text) for p in d.paragraphs)}")
print(f"  Headings (all levels): {len(heading_paras)}")
print(f"  TOC entries: {len(toc_table.rows)}")
print(f"  File size: {os.path.getsize(FILE)//1024} KB")
