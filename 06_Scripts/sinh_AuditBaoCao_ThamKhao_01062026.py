# -*- coding: utf-8 -*-
"""Sinh bao cao audit chi tiet cua file ThamKhao_Titles_TiengViet_v2.
Liet ke tat ca loi phat hien khi kiem tra tung tu, tung cau.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'AuditBaoCao_ThamKhao_Titles_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.4


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(text, italic=False, indent=True):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def WARN(text):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_table(headers, rows, col_widths_cm):
    t = d.add_table(rows=1, cols=len(headers)); t.style = 'Light Grid Accent 1'
    t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


H1('BÁO CÁO AUDIT CHI TIẾT')
P('File audit: ThamKhao_Titles_Q1Q3_AsiaChauPhi_TiengViet_v2_01062026.docx',
  italic=True, indent=False)
P('Phương pháp: kiểm tra từng từ, từng câu, từng paper bằng search web '
  '(PubMed/PMC + nature.com + frontiers + scimago + JCR)', italic=True, indent=False)
P('Ngày soạn: 01/06/2026', italic=True, indent=False)


# ============================================================
H1('TÓM TẮT PHẠM VI AUDIT')

P('Em đã kiểm tra TỪNG entry trong file v2 — tổng cộng 33 paper references + '
  '10 đề xuất tiêu đề + 4 dạng pattern Q1 + 3 dạng pattern Q3 + 10 chỉ số '
  'tạp chí. Quá trình audit phát hiện một số lỗi nghiêm trọng cần sửa '
  'trước khi gửi file cho đồng tác giả Nguyễn Minh Đức.')

P('Báo cáo chia thành 3 loại lỗi: (1) LỖI VỀ TÁC GIẢ — em ghi sai tên tác '
  'giả đầu hoặc nhầm sample; (2) LỖI VỀ NĂM XUẤT BẢN — em ghi sai năm thực '
  'sự; (3) LỖI VỀ ĐỐI TƯỢNG NGHIÊN CỨU — paper không phải về thanh thiếu '
  'niên như em ghi.')


# ============================================================
H1('1. LỖI VỀ TÁC GIẢ + SAMPLE — 3 paper')

errors_author = [
    ('Paper #1 — Risk factors of depressive and anxiety symptoms in '
     'Chinese adolescents',
     'Em ghi: "Wang và cộng sự — Trung Quốc"',
     'Đúng: **Juan J, Li J, Wang X et al.** — Wang chỉ là tác giả thứ 3, '
     'không phải đầu. Title gốc còn có "...in Chinese adolescent GIRLS" — '
     'em ghi sót "girls" (mẫu chỉ là nữ giới)',
     'Scientific Reports 2025; DOI: 10.1038/s41598-025-16396-5'),
    ('Paper #17 — Rwandese GAD-7 measurement invariance',
     'Em ghi: "Tác giả Trung Quốc (lab Rwanda)"',
     'Đúng: **Lisa Cynthia Niwenahisemo, Su Hong, Li Kuang** (từ Chongqing '
     'Medical University, study trên 1,813 HS Kigali, Rwanda). Em ghi "lab '
     'Rwanda" không chuẩn xác — họ là Chinese researchers nghiên cứu mẫu '
     'Rwanda',
     'Frontiers in Psychiatry 2024; DOI: 10.3389/fpsyt.2024.1346267'),
    ('Paper #9 — Hanoi Vietnam internet addiction',
     'Em ghi: "Tác giả Việt Nam (Hà Nội)" + "PMC\n2023"',
     'Đúng: **Tran Minh Dien, Pham Thi Lan Chi, Pham Quang Duy, Le Ha Anh, '
     'Nguyen Thi Kim Ngan, Vu Thi Hoang Lan**. n=5,325 HS THCS Hà Nội. '
     'Journal cụ thể là **BMC Public Health** (em ghi chung "PMC")',
     'BMC Public Health 2023; DOI: 10.1186/s12889-023-17348-2'),
]

make_table(['Paper', 'Em ghi (SAI)', 'Đúng (sau verify)', 'Nguồn'],
           errors_author,
           [5.0, 4.0, 6.0, 3.0])


# ============================================================
H1('2. LỖI VỀ NĂM XUẤT BẢN — 2 paper')

errors_year = [
    ('Paper #13 — Ethiopia non-recursive SEM',
     'Em ghi: 2023',
     'Đúng: **2024** (10/04/2024). Em nhầm với năm submit (2/2023) — paper '
     'thực sự published April 2024',
     'PLOS ONE 2024; DOI: 10.1371/journal.pone.0281571'),
    ('Paper #6 — Personality traits Chinese adolescents Frontiers Psychology',
     'Em ghi: 2025',
     'Đúng: **2026** (16/01/2026). Paper đã DOI 2025 nhưng published year '
     'là 2026',
     'Frontiers in Psychology 2026; DOI: 10.3389/fpsyg.2025.1748370'),
]

make_table(['Paper', 'Em ghi (SAI)', 'Đúng (sau verify)', 'Nguồn'],
           errors_year,
           [5.0, 3.0, 6.0, 4.0])


# ============================================================
H1('3. LỖI VỀ ĐỐI TƯỢNG NGHIÊN CỨU — 1 paper')

WARN('NGHIÊM TRỌNG: Paper #7 không phải về thanh thiếu niên — em đã đưa '
     'nhầm vào danh sách paper adolescent.')

errors_subject = [
    ('Paper #7 — Social support and anxiety moderated mediating model',
     'Em ghi: dùng làm ví dụ Q1 cho "thanh thiếu niên"',
     'Đúng: Mẫu là **1,097 SINH VIÊN ĐẠI HỌC** từ Hồ Nam (Hunan), Trung '
     'Quốc — KHÔNG phải adolescent (thanh thiếu niên). Em đã đưa nhầm '
     'vào danh sách paper về adolescent → nhóm sẽ hiểu sai về nhóm '
     'mẫu nếu dùng làm tham khảo',
     'Scientific Reports 2025; DOI: 10.1038/s41598-025-14336-x'),
]

make_table(['Paper', 'Em ghi (SAI)', 'Đúng (sau verify)', 'Nguồn'],
           errors_subject,
           [5.0, 4.0, 6.0, 3.0])

P('Khuyến nghị: **Bỏ paper #7 khỏi danh sách**, hoặc thay bằng paper khác '
  'thực sự về adolescent. Em đã search nhanh và thấy có nhiều paper '
  'mediation alternative phù hợp hơn.', italic=True)


# ============================================================
H1('4. ĐÁNH GIÁ CHẤT LƯỢNG TỔNG THỂ')

H2('Số lượng lỗi đã phát hiện')

make_table(
    ['Loại lỗi', 'Số paper bị ảnh hưởng', 'Tỷ lệ'],
    [
        ('Lỗi tác giả/sample (Nhóm 1)', '3', '9,1%'),
        ('Lỗi năm xuất bản (Nhóm 2)', '2', '6,1%'),
        ('Lỗi đối tượng nghiên cứu (Nhóm 3)', '1', '3,0%'),
        ('Lỗi IF + acceptance rate (đã sửa trong v2)', '9 tạp chí', '—'),
        ('Tổng số paper có lỗi', '6/33', '18,2%'),
    ],
    [6.0, 5.5, 5.5]
)

P('Trên tổng 33 paper được liệt kê, em phát hiện 6 paper có vấn đề (18%). '
  'Đây là tỷ lệ cao hơn em mong đợi và cho thấy em cần thận trọng hơn khi '
  'compile bộ tham khảo nhanh — nên search verify từng entry trước khi '
  'đưa vào danh sách.', italic=True)

H2('Em xin nhận lỗi')

P('Em xin lỗi vì đã không verify kỹ từng entry trong các phiên trước, dẫn '
  'đến nhiều lỗi đã đi vào file gửi đến thầy và nhóm. Bài học rút ra: '
  'với loại tài liệu tham khảo có tính dữ liệu cụ thể như tên paper, năm, '
  'tác giả, tạp chí, em phải LUÔN verify từng entry bằng search trực tiếp '
  'trên PubMed/Crossref trước khi đưa vào file gửi đi.')


# ============================================================
H1('5. ĐỀ XUẤT HÀNH ĐỘNG')

P('Em đề xuất ba phương án để xử lý các lỗi này, tùy theo mức độ ưu tiên '
  'mà thầy quyết định:', indent=False)

H2('Phương án A: Em rebuild v3 với tất cả các sửa')

P('Em sẽ tạo file v3 với:', indent=False)
P('  • Sửa tác giả Paper #1 (Juan J et al.) + thêm "GIRLS" vào title', indent=False)
P('  • Sửa năm Paper #13 (2023→2024)', indent=False)
P('  • Sửa năm Paper #6 (2025→2026)', indent=False)
P('  • Bỏ hoặc thay Paper #7 (vì là sinh viên ĐH, không phải adolescent)', indent=False)
P('  • Bổ sung thông tin đầy đủ tác giả + journal cho Paper #9 (Tran Minh '
  'Dien et al., BMC Public Health)', indent=False)
P('  • Bổ sung tác giả cho Paper #13 (Gebreegziabher Z et al.) và Paper #17 '
  '(Niwenahisemo LC et al.)', indent=False)
P('  • Bổ sung sample sizes cho Paper #5 (n=486), Paper #6 (n=3,673), '
  'Paper #17 (n=1,813)', indent=False)

P('Thời gian: ~15-20 phút.', italic=True, indent=False)


H2('Phương án B: Em chỉ ghi chú các sửa vào file v2 hiện tại')

P('Em sẽ thêm phần "ERRATA — Các sửa đã phát hiện" ở cuối file v2, không '
  'tạo file mới. Thời gian: ~5 phút.', indent=False)


H2('Phương án C: Để các lỗi nguyên trong file v2 (không khuyến nghị)')

P('Phương án này CHỈ chấp nhận được nếu thầy không gửi file v2 cho đồng '
  'tác giả Nguyễn Minh Đức. Em không khuyến nghị vì các lỗi tác giả + năm '
  'sẽ gây nhầm lẫn cho người đọc khi muốn tra cứu paper.', indent=False)


P('Em đề xuất **Phương án A** vì các lỗi đủ nghiêm trọng để cần fix '
  'trước khi file đi đến đồng tác giả. Thầy xác nhận em sẽ thực hiện ngay.',
  italic=True, indent=False)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
