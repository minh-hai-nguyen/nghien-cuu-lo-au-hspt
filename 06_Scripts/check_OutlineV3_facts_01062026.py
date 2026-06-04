# -*- coding: utf-8 -*-
"""B4: Verify all citations + facts in outline v3 (Q1 + Q3) vs PDF catalog + LA.
01/06/2026."""
import os, sys, io, json, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Read both v3 outlines
files = {
    'Q1_v3': os.path.join(ROOT, 'bai-bao-Q1', 'Outline_Q1_v3_01062026.docx'),
    'Q3_v3': os.path.join(ROOT, 'bai-bao-Q1', 'Outline_Q3_v3_01062026.docx'),
}

texts = {}
for k, p in files.items():
    d = Document(p)
    t = '\n'.join(para.text for para in d.paragraphs)
    for tb in d.tables:
        for row in tb.rows:
            for c in row.cells:
                t += '\n' + c.text
    texts[k] = t

# Load PDF catalog
with open(os.path.join(ROOT, '06_Scripts', 'pdf_catalog_29052026.json'),
          'r', encoding='utf-8') as f:
    cat = json.load(f)

# Citations to verify
citations = {
    'Xu': '2021', 'Chen': '2023', 'Wen': '2020', 'Anderson': '2025',
    'Compas': '2017', 'Pascoe': '2020', 'Steare': '2023',
    'Jefferies': '2020', 'Nakie': '2022', 'V-NAMHS': '2022',
    'Hoang Trung Hoc': '2025', 'Bhardwaj': '2020', 'Qiu': '2022',
    'Zhu': '2025', 'Alharbi': '2019', 'Saikia': '2023',
    'Chorpita': '2000', 'Sun': '2011', 'Kwon': '2013', 'Olweus': '1996',
    'Goodenow': '1993', 'Zimet': '1988', 'Rosenberg': '1965',
    'Lazarus & Folkman': '1984', 'Carver': '1997',
    'Hu & Bentler': '1999', 'Cheung & Rensvold': '2002',
    'Braun & Clarke': '2006', 'Creswell & Plano Clark': '2018',
    'McLean': '2011', 'Bronfenbrenner': '2006',
    'Triandis': '1995', 'Markus & Kitayama': '1991', 'Allen': '2013',
    'Beesdo-Baum': '2022', 'Nunnally': '1978',
    'GBD ASEAN': '2025',
}

# Facts
key_facts = [
    '1.352', '1,352', '614', '738',
    '55.82', '55,82', '25.06', '25,06', '48.41', '48,41',
    '44.484', '44,484', '45.984', '45,984', '0.246', '0,246',
    '64.28', '64,28', '27.88', '27,88', '56.98', '56,98',
    '32.13', '32,13', '19.46', '19,46',
    'BMC Psychiatry', 'PLOS ONE', '+0.510', '+0.533', '-0.457',
    '+0.376',
    'Hang Thi Cong', 'Nguyen Minh Duc', 'Duc Minh Dao',
]

print('='*70)
print('B4: VERIFY OUTLINE v3 (Q1 + Q3) FACTS')
print('='*70)
print()

# Check citations
print('=== CITATIONS in Q1 v3 + Q3 v3 ===')
for c, y in citations.items():
    q1 = c in texts['Q1_v3']
    q3 = c in texts['Q3_v3']
    # Check PDF exists in catalog
    pdf_match = any(c.lower() in pdf['filename'].lower() or
                    c.lower().replace(' ', '_') in pdf['filename'].lower()
                    for pdf in cat['pdfs'])
    in_q1 = 'Q1✓' if q1 else 'Q1-'
    in_q3 = 'Q3✓' if q3 else 'Q3-'
    has_pdf = 'PDF✓' if pdf_match else 'PDF?'
    print(f'  {in_q1} {in_q3} {has_pdf} | {c} ({y})')

print()
print('=== KEY FACTS in Q1 v3 + Q3 v3 ===')
for fact in key_facts:
    q1 = fact in texts['Q1_v3']
    q3 = fact in texts['Q3_v3']
    if q1 or q3:
        in_q1 = 'Q1✓' if q1 else 'Q1-'
        in_q3 = 'Q3✓' if q3 else 'Q3-'
        print(f'  {in_q1} {in_q3} | {fact}')

# Special checks: BLOCKING placeholders
print()
print('=== BLOCKING PLACEHOLDERS ===')
for marker in ['[NCS confirm]', '[NCS/Thầy confirm]', 'BLOCKING']:
    q1_count = texts['Q1_v3'].count(marker)
    q3_count = texts['Q3_v3'].count(marker)
    print(f'  Q1: {q1_count}× | Q3: {q3_count}× | "{marker}"')

# Count overall
print()
print('='*70)
print('SUMMARY')
print('='*70)
cit_q1 = sum(1 for c in citations if c in texts['Q1_v3'])
cit_q3 = sum(1 for c in citations if c in texts['Q3_v3'])
print(f'Q1 v3: {cit_q1}/{len(citations)} citations present')
print(f'Q3 v3: {cit_q3}/{len(citations)} citations present')
print(f'Word count Q1: ~{len(texts["Q1_v3"].split())}')
print(f'Word count Q3: ~{len(texts["Q3_v3"].split())}')
