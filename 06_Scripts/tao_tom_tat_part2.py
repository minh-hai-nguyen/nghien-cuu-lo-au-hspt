# -*- coding: utf-8 -*-
"""Tao tom tat bai 44-58 (15 bai con lai)"""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)

def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')):
        tcW.remove(old)
    tcW.append(w)

def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')):
            tblPr.remove(old)
        tblPr.append(layout)
    tblGrid = t._tbl.find(qn('w:tblGrid'))
    if tblGrid is not None:
        t._tbl.remove(tblGrid)
    tblGrid = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol')
        gc.set(qn('w:w'), str(int(w * 567)))
        tblGrid.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tblGrid)

def make_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3); s.right_margin = Cm(2)
    return doc

def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    return p

def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
    return t

def save(doc, name):
    p = os.path.join(OUT_DIR, name)
    doc.save(p)
    d = Document(p)
    chars = sum(len(x.text) for x in d.paragraphs)
    print(f'  {name}: {chars} chars, {len(d.tables)} tables')

# ============================================================
# BÀI 44 — Walder 2025 JMIR DMHI SAD MA
# ============================================================
print('[6/20] Bài 44: Walder 2025 JMIR DMHI SAD')
doc = make_doc()
H(doc, 'Tóm tắt bài QT40 — Meta-analysis Can thiệp Sức khoẻ Tâm thần Số (DMHI) cho lo âu xã hội', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Digital mental health interventions for prevention and treatment of social anxiety disorder in children, '
       'adolescents and young adults: Systematic review and meta-analysis" của Walder, N., Frey, A., Berger, T. và cộng sự '
       '(2025), đăng trên Journal of Medical Internet Research (JMIR — Q1, IF ≈ 7,4). DOI 10.2196/preprints.67067. PROSPERO '
       'CRD42023424181. Phân tích tổng hợp 21 RCT về DMHI cho rối loạn lo âu xã hội (SAD) ở trẻ em, vị thành niên và '
       'thanh niên (TB < 25 tuổi).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống RCT đánh giá DMHI (internet, app, VR) cho lo âu xã hội ở người trẻ. 2 tác giả độc lập sàng lọc, '
       'trích xuất, đánh giá thiên lệch. Random-effects meta-analysis sử dụng Hedges g. Phân tích phụ theo: nền tảng (CBT '
       'vs khác), thiết kế (SAD-specific vs tổng quát), có hướng dẫn người vs tự học, đối chứng (WL vs khác).')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'DMHI có hiệu lực TRUNG BÌNH so với bất kỳ đối chứng nào (g = 0,508; KTC 95 %: 0,308–0,707). Hai phát hiện cốt lõi: '
       '(1) DMHI thiết kế ĐẶC THÙ cho SAD có hiệu lực LỚN NHẤT (g = 0,878), gần gấp đôi DMHI tổng quát; (2) DMHI có '
       'HƯỚNG DẪN người (g = 0,825) hiệu lực gần gấp đôi DMHI không hướng dẫn (g ≈ 0,3–0,4).')

table(doc, ['Phân nhóm', 'Hedges g', 'KTC 95 %', 'Ý nghĩa'],
      [['DMHI vs đối chứng tổng', '0,508', '0,308–0,707', 'TRUNG BÌNH'],
       ['DMHI vs danh sách chờ', '0,576', '0,343–0,809', 'Mạnh hơn vs WL'],
       ['DMHI nền CBT', '0,610', '0,361–0,859', 'CBT nền tảng tốt'],
       ['DMHI có hướng dẫn', '0,825', '0,425–1,224', 'LỚN — quan trọng'],
       ['DMHI thiết kế SAD-specific', '0,878', '0,469–1,278', 'LỚN NHẤT'],
       ['DMHI không hướng dẫn', '~ 0,3–0,4', '—', 'Nhỏ — kém hiệu quả']],
      widths=[5.5, 2.5, 3.5, 3.5])

H(doc, 'Phản biện', level=2)
P(doc, 'Đây là bằng chứng MẠNH NHẤT về DMHI / iCBT cho SAD ở trẻ và VTN trong toàn bộ tài liệu. JMIR Q1 IF ≈ 7,4. PROSPERO '
       'tiền đăng ký. Two-key takeaways có ý nghĩa thực hành rõ ràng: app phải SAD-specific + có hỗ trợ người. Hạn chế: '
       'chủ yếu các RCT từ phương Tây + Đông Á phát triển; ít dữ liệu LMIC; định nghĩa "guided" rất khác nhau giữa các '
       'nghiên cứu (chat, email, video); chưa rõ liều tối ưu (số tuần, số module).')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Phát triển app iCBT tiếng Việt cho SAD với hỗ trợ chat người, dựa trên hai điều kiện cốt lõi của Walder. So sánh '
       'mức độ guided khác nhau (chat thường vs định kỳ vs sự kiện). RCT VN cho HS THCS/THPT với app này.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao. MA 21 RCT, JMIR Q1, định hướng phát triển app.', bold=True)
save(doc, 'QT40_Walder_DMHI_SAD_2025.docx')

# ============================================================
# BÀI 45 — Hải Phòng 2024 (DASS-21 420 HS THPT)
# ============================================================
print('[7/20] Bài 45: Hải Phòng 2024')
doc = make_doc()
H(doc, 'Tóm tắt bài VN25 — Sức khoẻ tâm thần học sinh THPT tại Hải Phòng', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Thực trạng sức khoẻ tâm thần và một số yếu tố liên quan ở học sinh trung học phổ thông tại Hải Phòng" '
       '(2023, đăng 2024), đăng trên Tạp chí Y học Dự phòng Việt Nam (VJPM). Mẫu nghiên cứu là 420 học sinh THPT tại '
       'Hải Phòng, sử dụng thang DASS-21 (Depression Anxiety Stress Scale 21 mục).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Cắt ngang trên 420 học sinh 4 trường THPT Hải Phòng. DASS-21 phiên bản tiếng Việt — đo cả 3 trục: trầm cảm, lo âu, '
       'căng thẳng. Bảng hỏi nhân khẩu học. Phân tích chi-square + hồi quy logistic đa biến. Đã được phê duyệt đạo đức.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ lo âu (DASS-21 ≥ 8) là 39,3 % — cao đáng kể. Trầm cảm và stress cũng có tỷ lệ cao tương tự. Nữ có tỷ lệ '
       'cao hơn nam (p < 0,05). Các yếu tố liên quan có ý nghĩa: áp lực học tập, mâu thuẫn gia đình, sử dụng MXH > 4h/ngày, '
       'thiếu ngủ < 6h, không có hoạt động thể chất.')

H(doc, 'Phản biện', level=2)
P(doc, 'Mẫu vừa phải (n = 420), 1 thành phố cảng phía Bắc — đại diện vùng đô thị ven biển. Sử dụng DASS-21 đo cả 3 trục — '
       'tốt cho so sánh với Dong 2025 (TQ) và các nghiên cứu khác. Tỷ lệ 39,3 % phù hợp với "baseline" lo âu VTN châu Á '
       '(40–55 %). Hạn chế: cắt ngang, mẫu thuận tiện, không có nhóm chứng, không phân tích sâu cơ chế.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'So sánh đa vùng (Hải Phòng vs Hà Nội vs TPHCM) bằng cùng công cụ DASS-21. Nghiên cứu dọc theo dõi cùng nhóm 2-3 '
       'năm. Phân tích trung gian (mediation) MXH → giấc ngủ → lo âu.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao cho dữ liệu vùng. Tạp chí trong nước nhưng đầy đủ.', bold=True)
save(doc, 'VN25_HaiPhong_2024.docx')

# ============================================================
# BÀI 46 — Zheng 2025 (TQ MXH+Stress+Sleep, mediation)
# ============================================================
print('[8/20] Bài 46: Zheng 2025 TQ MXH-Stress-Sleep')
doc = make_doc()
H(doc, 'Tóm tắt bài QT41 — Tác động MXH, căng thẳng học tập, giấc ngủ lên lo âu (mô hình trung gian)', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"The Effects of Social Media Addiction, Academic Stress, and Sleep Quality on Anxiety Symptoms: A Cross-Sectional '
       'Study of Chinese Vocational Students" của Zheng, G.F. & Peng, H.Y. (2025), đăng trên Psychology Research and '
       'Behavior Management vol. 18, tr. 1571–1584. DOI 10.2147/PRBM.S522652. Open Access. Mẫu: 469 học sinh trường nghề '
       '12–18 tuổi tại Quảng Đông, Trung Quốc.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Cắt ngang với phân tích trung gian (mediation analysis). Công cụ: SMAS (Social Media Addiction Scale 6 mục), '
       'Academic Stress Scale 16 mục, PSQI (Pittsburgh Sleep Quality Index), GAD-7 (lo âu), GSES (General Self-Efficacy '
       'Scale). Mô hình tam giác: MXH + áp lực HT + giấc ngủ → triệu chứng lo âu, với self-efficacy là biến TRUNG GIAN.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tương quan trực tiếp với lo âu: giấc ngủ kém r = 0,816 (MẠNH NHẤT), áp lực HT r = 0,505, MXH r = 0,415. Hồi quy '
       'đa biến: giấc ngủ β = 0,615 (mạnh nhất), áp lực β = 0,223, MXH β = 0,153. PHÁT HIỆN QUAN TRỌNG: self-efficacy là '
       'biến TRUNG GIAN — MXH gián tiếp 63,13 % tổng tác động (chủ yếu qua self-efficacy, không trực tiếp); áp lực 55,84 %; '
       'giấc ngủ 24,63 %.')

table(doc, ['Yếu tố', 'r với lo âu', 'β hồi quy', '% gián tiếp qua self-efficacy'],
      [['Giấc ngủ kém (PSQI)', '0,816', '0,615', '24,63 %'],
       ['Áp lực học tập', '0,505', '0,223', '55,84 %'],
       ['MXH (SMAS)', '0,415', '0,153', '63,13 %']],
      widths=[5.0, 2.5, 2.5, 5.5])

H(doc, 'Phản biện', level=2)
P(doc, 'Mô hình tam giác đo CẢ 3 yếu tố — ít NC nào làm được. Phân tích trung gian phát hiện CƠ CHẾ rõ ràng: MXH ảnh '
       'hưởng lo âu CHỦ YẾU qua giảm self-efficacy, không trực tiếp. Phù hợp Zhu QT05 (giấc ngủ < 5h AOR = 13,71). GAD-7 '
       'cùng công cụ với Hoa 2024 và A8 VN COVID. Hạn chế: HS nghề (ít được nghiên cứu), 1 tỉnh, cắt ngang.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Lặp lại với HS phổ thông VN. Nghiên cứu dọc kiểm tra hướng nhân quả. Can thiệp tăng self-efficacy như đường dẫn '
       'để giảm tác động xấu của MXH.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Mediation rõ ràng, phát hiện cơ chế quan trọng.', bold=True)
save(doc, 'QT41_Zheng_MXH_Stress_Sleep_2025.docx')

# ============================================================
# BÀI 47 — Long An 2025
# ============================================================
print('[9/20] Bài 47: Long An 2025')
doc = make_doc()
H(doc, 'Tóm tắt bài VN26 — Lo âu ở học sinh THPT tại Long An', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình về thực trạng lo âu ở học sinh THPT tại tỉnh Long An (Đồng bằng sông Cửu Long), đăng trên Tạp chí Y '
       'học Việt Nam (VJOL) năm 2025. Khảo sát học sinh THPT tại Long An.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Cắt ngang trên học sinh THPT các trường tại Long An. Sử dụng thang đo lo âu (DASS-21 hoặc tương đương) phiên bản '
       'tiếng Việt. Bảng hỏi nhân khẩu học và yếu tố liên quan.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tỷ lệ lo âu là 57,2 % — CAO NHẤT trong số các nghiên cứu lo âu HS THPT tại Việt Nam được hệ thống tổng hợp đến '
       'thời điểm hiện tại. Tỷ lệ này cao hơn rõ rệt so với Hà Nội (Hoa 2024: 40,6 %), Hải Phòng (39,3 %), Huế (Trần Thảo '
       'Vi 2024: ~ 30–40 %). Yếu tố nguy cơ: nữ giới, áp lực thi đại học, mâu thuẫn gia đình, lạm dụng MXH.')

H(doc, 'Phản biện', level=2)
P(doc, 'Tỷ lệ 57,2 % là CAO BẤT THƯỜNG so với các vùng khác — cần giải thích. Có thể do: (1) cut-off thấp; (2) đặc thù '
       'vùng (Long An giáp TPHCM, áp lực thi vào ĐH cao); (3) thiên lệch chọn mẫu; (4) ảnh hưởng còn lại của COVID-19. '
       'Tạp chí trong nước, mẫu cắt ngang, không có nhóm chứng. Cần đối chiếu với phương pháp đo cụ thể.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'So sánh trực tiếp với các tỉnh ĐBSCL khác bằng cùng công cụ. Phân tích yếu tố vùng đặc thù. Cần điều tra lý do '
       'tỷ lệ cao bất thường: định lượng + định tính.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐ Trung bình. Có giá trị dữ liệu vùng, nhưng cần cẩn trọng diễn giải tỷ lệ cao.', bold=True)
save(doc, 'VN26_LongAn_2025.docx')

# ============================================================
# BÀI 48 — Brown & Carter 2025 (UK Editorial)
# ============================================================
print('[10/20] Bài 48: Brown & Carter 2025 UK Editorial')
doc = make_doc()
H(doc, 'Tóm tắt bài QT42 — Editorial UK về can thiệp tại trường cho trầm cảm và lo âu', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"School-based interventions for depression and anxiety in UK" của Brown, J.S.L. & Carter, B. (2025), đăng trên '
       'Journal of Mental Health vol. 34(4), tr. 357–361 (Q1, IF ≈ 3,5). DOI 10.1080/09638237.2025.2512332. Đây là '
       'editorial từ chuyên gia hàng đầu UK về can thiệp tâm lý cộng đồng — June S.L. Brown.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Editorial / commentary tổng hợp bằng chứng từ nhiều RCT và mô hình can thiệp tại trường ở Anh quốc. Phân loại 4 '
       'mô hình: phổ quát mindfulness, BESST CBT self-referral, MHST (Mental Health Support Teams), và PLACES self-referral. '
       'So sánh hiệu quả + bài học triển khai.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tổng hợp 4 mô hình can thiệp tại trường UK với các phát hiện then chốt:')
table(doc, ['Mô hình', 'Người cung cấp', 'Hiệu quả', 'Bằng chứng'],
      [['Mindfulness phổ quát', 'Giáo viên', 'THẤT BẠI', '8.376 HS / 85 trường (Kuyken 2022)'],
       ['BESST (CBT self-referral)', 'Chuyên gia tâm lý', 'DƯƠNG TÍNH', '900 HS / 57 trường (Brown 2024)'],
       ['MHST (Mental Health Support Team)', 'Thạc sĩ trị liệu', 'HỨA HẸN', 'Mới triển khai NHS UK'],
       ['PLACES (self-referral)', 'Đa cấp', 'HỨA HẸN', 'Dùng từ thường ngày, giảm kỳ thị']],
      widths=[5.0, 3.0, 2.5, 5.0])

P(doc, 'Năm bài học từ Editorial: (1) Universal MINDFULNESS thất bại do engagement thấp ở HS; (2) TARGETED CBT (HS có '
       'triệu chứng tự đăng ký) thành công; (3) Chuyên gia lâm sàng > GV về hiệu quả, nhưng cả hai đều khả thi; (4) Co-design '
       '(cho HS tham gia thiết kế) tăng engagement; (5) Ngôn ngữ THƯỜNG NGÀY ("căng thẳng" thay "trầm cảm") giảm kỳ thị.')

H(doc, 'Phản biện', level=2)
P(doc, 'Editorial — không phải SR/MA gốc. Ý kiến chuyên gia + tổng hợp bằng chứng. Tuy nhiên cung cấp khung phân loại '
       'rất hữu ích. Cảnh báo về universal interventions là QUAN TRỌNG cho thiết kế đề cương Việt Nam. Phù hợp với Zhang, '
       'Liang & Kang 2026 (Bayesian MA — universal CBT hiệu ứng nhỏ). Hạn chế: chỉ UK (NHS context), khác bối cảnh LMIC.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Áp dụng mô hình BESST (targeted + self-referral) cho VN — RCT thử nghiệm. Adapt PLACES language giảm kỳ thị tiếng '
       'Việt. So sánh chuyên gia vs GV cung cấp CBT trong bối cảnh VN.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Editorial Q1, khung phân loại + bài học rõ ràng cho thiết kế.', bold=True)
save(doc, 'QT42_Brown_Carter_UK_2025.docx')

# ============================================================
# BÀI 49 — Bress 2024 JAMA Maya app
# ============================================================
print('[11/20] Bài 49: Bress 2024 JAMA Maya')
doc = make_doc()
H(doc, 'Tóm tắt bài QT43 — RCT ứng dụng Maya (CBT-app) cho lo âu VTN', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Efficacy of a Mobile App-Based Intervention for Young Adults With Anxiety Disorders: A Randomized Clinical Trial" '
       'của Bress, J.N., Falk, A., Schier, M.M. và cộng sự (2024), đăng trên JAMA Network Open vol. 7(8), e2428372 (Q1, '
       'IF ≈ 13,8). DOI 10.1001/jamanetworkopen.2024.28372. ClinicalTrials.gov NCT05130281. Đối tượng: 59 thanh niên 18–25 '
       'tuổi có rối loạn lo âu (GAD 56 %, SAD 41 %), tuyển qua quảng cáo trực tuyến.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'RCT 3 nhánh khuyến khích dựa trên tin nhắn (gain-frame, loss-frame, social support). Can thiệp: ứng dụng Maya — '
       'CBT tự hướng dẫn 12 phiên / 6 tuần. Đo lường: HAM-A (chính), ASI, LSAS, HDRS-24, uMARS. Theo dõi đến tuần 12. '
       'Phân tích linear mixed-effects model, intention-to-treat (ITT). Tuân thủ CONSORT.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Lo âu giảm có ý nghĩa thống kê và HIỆU LỰC TĂNG theo thời gian:')

table(doc, ['Thời điểm', 'Chênh lệch HAM-A', 'KTC 95 %', "Cohen's d", 'p'],
      [['Tuần 3 (giữa)', '−3,20', '−4,76 đến −1,64', '0,64 (TB)', '< 0,001'],
       ['Tuần 6 (kết thúc)', '−5,64', '−7,23 đến −4,05', '0,94 (LỚN)', '< 0,001'],
       ['Tuần 12 (theo dõi)', '−5,67', '−7,29 đến −4,04', '1,04 (RẤT LỚN)', '< 0,001']],
      widths=[3.5, 3.2, 3.5, 3.0, 1.8])

P(doc, 'Trên các thang đo phụ: ASI (anxiety sensitivity) Cohen d = 0,93 (tuần 6); LSAS (lo âu xã hội) d = 1,07 (tuần 12). '
       'Giả thuyết chính (gain-frame + social support hiệu quả nhất) KHÔNG được ủng hộ — cả 3 nhánh đều cho hiệu quả tương '
       'đương, gợi ý rằng chính bản thân app CBT là yếu tố then chốt.')

H(doc, 'Phản biện', level=2)
P(doc, 'JAMA Network Open Q1 IF ≈ 13,8 — bằng chứng chất lượng cực cao. RCT đăng ký pre-registered. Hiệu lực Cohen d = '
       '1,04 là RẤT LỚN cho can thiệp app — gần như tương đương CBT mặt đối mặt. Hạn chế: (1) đối tượng 18–25 tuổi (THANH '
       'NIÊN), không phải VTN học sinh — cần kiểm chứng ở tuổi nhỏ hơn; (2) tuyển mộ qua quảng cáo trực tuyến → mẫu tự '
       'chọn cao, có động lực sẵn → kết quả có thể không generalizable cho dân số chung; (3) cỡ mẫu n = 59 không quá lớn.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Adapt Maya cho HS THCS/THPT với UI/UX phù hợp tuổi nhỏ. RCT VN cho VTN cấp 3 + sinh viên năm 1. Phát triển phiên '
       'bản tiếng Việt với hỗ trợ chat người (theo Walder 2025).')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐⭐ Rất cao. RCT JAMA Q1, hiệu lực Cohen d = 1,04 cho mobile app CBT.', bold=True)
save(doc, 'QT43_Bress_Maya_App_2024.docx')

print('11/20 done.')
