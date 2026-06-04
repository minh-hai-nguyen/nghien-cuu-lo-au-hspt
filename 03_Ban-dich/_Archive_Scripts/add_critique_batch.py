# -*- coding: utf-8 -*-
"""Thêm phần PHẢN BIỆN (CRITICAL REVIEW) cho các bài dịch cũ thiếu"""
import os, sys
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import add_red_heading, add_red, add_abbreviation_table
import docx

OUT = os.path.dirname(os.path.abspath(__file__))

# Phản biện cho từng bài — trích dẫn dày
critiques = {
    '02_Hoa_2024_Frontiers.docx': {
        'strengths': [
            'Frontiers in Public Health Q1, PubMed. GAD-7 alpha=0,916 — rất cao.',
            'Phương pháp hỗn hợp (mixed methods) sáng tạo: định lượng (GAD-7) + phỏng vấn sâu — ít NC VN làm (Pham 2024 cũng dùng mixed methods).',
            'n = 3.910 — mẫu lớn nhất trong các NC lo âu VTN VN (so: Bảo Quyên 2025 n=501, Vĩnh Lộc 2024 n=976, Danh Lâm 2022 n=482).',
            'Xác nhận nữ > nam (p < 0,01) — phù hợp xu hướng toàn cầu: 59 Countries AOR=1,51 (Islam et al. 2025), Ireland 2024 nữ tăng nhanh hơn.',
        ],
        'limits': [
            'Chỉ Hà Nội (đô thị) — không đại diện nông thôn/DTTS. So: Ngô Anh Vinh 2024 (VN15) DTTS Lạng Sơn 54,4%.',
            'Cắt ngang — không suy luận nhân quả. Li 2025 (QT22) dùng thiết kế dọc cho screen time.',
            'GAD-7 sàng lọc: 40,6% ≠ 40,6% có rối loạn lo âu. So V-NAMHS 2022 (DISC-5 chẩn đoán: 2,3%) — chênh 17 lần.',
            'Phỏng vấn sâu: không nêu rõ số lượng (20 HS) và cách chọn mẫu định tính.',
            'Bối cảnh hậu COVID (10-11/2021) — tỷ lệ có thể cao hơn bình thường. Hoàng Trung Học 2025: lo âu giảm từ 41,5% xuống 25,4% sau COVID.',
        ],
        'gaps': [
            'Mở rộng ra nông thôn và vùng DTTS — so sánh với Ngô Anh Vinh 2024 (DTTS 54,4%), An Giang 2025 (nông thôn 61,2%).',
            'So sánh GAD-7 với DISC-5 trên cùng mẫu để ước tính khoảng cách sàng lọc–chẩn đoán (Gap #2 cross-study).',
            'NC dọc theo dõi lo âu hậu COVID — Hoàng Trung Học có 2 thời điểm nhưng cần nhiều hơn.',
            'Thêm biến screen time (Chen 2023: game OR=5,00; Norway 2025 QT21: MXH giải thích xu hướng).',
        ],
    },
    '03_GBD_ASEAN_2025_Lancet.docx': {
        'strengths': [
            'Lancet — tạp chí y khoa hàng đầu thế giới (IF > 160).',
            'GBD 2021 — dữ liệu toàn diện nhất, 204 quốc gia/vùng lãnh thổ.',
            'ASEAN 10 nước: 80,4 triệu ca RLTT — quy mô khổng lồ.',
            'VN đứng thứ 3 ASEAN về gánh nặng — số liệu quan trọng cho chính sách.',
        ],
        'limits': [
            'GBD dựa trên MÔ HÌNH ước tính — nhiều nước ĐNA thiếu dữ liệu gốc. So: Zhang et al. 2025 (QT30) cũng dùng GBD nhưng nhấn mạnh hạn chế SDI thấp.',
            'Không tách lo âu riêng — gộp chung rối loạn tâm thần. JAACAP 2024 (QT23) tách riêng: lo âu tăng gấp đôi.',
            'Sinh thái hồi cứu — xu hướng cấp quốc gia, không cá nhân.',
        ],
        'gaps': [
            'VN cần khảo sát quốc gia SKTT VTN riêng (không chỉ GBD ước tính). V-NAMHS 2022 là bước đầu.',
            'So sánh VN với Indonesia, Philippines, Thái Lan — cùng ASEAN nhưng tỷ lệ khác. Jefferies 2020 (QT35): VN 30,7% SAD vs Indonesia 22,9%.',
        ],
    },
    '04_Zhameden_2025_PLOSONE.docx': {
        'strengths': [
            'PLOS ONE Q1. Tổng quan hệ thống duy nhất về can thiệp SKTT trường học tại LMIC.',
            'Bao gồm 6 NC RCT/quasi từ LMIC — bối cảnh phù hợp VN.',
        ],
        'limits': [
            'Chỉ 3/4 NC hiệu quả cho trầm cảm, 1/4 cho lo âu — can thiệp kém hiệu quả cho lo âu.',
            '0 RCT từ VN hoặc ĐNA — GAP lớn nhất. BMC NMA 2025 (QT29): 30 RCT nhưng từ phương Tây.',
            'Thiếu đa dạng can thiệp — chủ yếu CBT. AJP 2024 (QT28): cần thêm ACT, PE, VRET.',
        ],
        'gaps': [
            'RCT can thiệp SKTT tại trường VN — ưu tiên #1 cross-study.',
            'CBT nhóm + PE kết hợp (BMC NMA 2025: cả hai hiệu quả) tại trường VN.',
        ],
    },
    '05_Anderson_2025_Wiley.docx': {
        'strengths': ['Tổng quan toàn diện yếu tố gia tăng lo âu VTN.', 'Wiley — nhà xuất bản uy tín.', 'Đa yếu tố: MXH, COVID, giáo dục, gia đình.'],
        'limits': ['Tổng quan narrative — không phải hệ thống/phân tích tổng hợp.', 'Chủ yếu phương Tây.'],
        'gaps': ['Cần tổng quan tương tự cho VN/ĐNA.', 'Decomposition (Norway 2025, QT21) — phương pháp mạnh hơn narrative review.'],
    },
    '06_VNAMHS_2022.docx': {
        'strengths': ['Khảo sát SKTT VTN quốc gia DUY NHẤT tại VN.', 'DISC-5/DSM-5 — chẩn đoán chuẩn hóa (không chỉ sàng lọc).', 'n = 5.996, 38 tỉnh.'],
        'limits': ['Chỉ 1 thời điểm — không có xu hướng. Norway (QT21): 13 năm, Korea (QT34): 16 năm.', 'Tỷ lệ 2,3% thấp hơn nhiều sàng lọc (40-86%) — cần so sánh trên cùng mẫu.', 'Tiếp cận dịch vụ chỉ 8,4% — thấp cả so với Philippines (2%, Puyat 2025).'],
        'gaps': ['Lặp lại V-NAMHS (3-5 năm/lần) để có xu hướng dài hạn.', 'So sánh DISC-5 vs GAD-7/DASS-21 trên cùng mẫu — gap #2 cross-study.'],
    },
    '07_Zhu_2025_BMC_China.docx': {
        'strengths': ['BMC Public Health Q1. n lớn (Suzhou, TQ).', 'Giấc ngủ <5h: AOR = 13,71 cho trầm cảm chắc chắn — yếu tố mạnh nhất.'],
        'limits': ['Chỉ TQ (Suzhou) — đô thị phát triển.', 'DASS-21 sàng lọc.'],
        'gaps': ['NC giấc ngủ-lo âu ở VTN VN — chưa có. 59 Countries (QT31): ngồi >4h AOR=1,50.'],
    },
    '08_Mudunna_2025_LancetSEA.docx': {
        'strengths': ['Lancet SEA. GBD ASEAN toàn diện.', '80,4 triệu ca — quy mô lớn.'],
        'limits': ['GBD ước tính — thiếu dữ liệu gốc nhiều nước ĐNA.'],
        'gaps': ['ASEAN cần khảo sát SKTT VTN chuẩn hóa (GSHS mở rộng).'],
    },
    '09_Puyat_2025_Filipino.docx': {
        'strengths': ['Dữ liệu Philippines — ĐNA, gần bối cảnh VN.', 'Trầm cảm tăng gấp đôi 2013-2021.'],
        'limits': ['Tiếp cận dịch vụ chỉ 2% — thấp nhất. So VN 8,4% (V-NAMHS).'],
        'gaps': ['So sánh VN vs Philippines — cùng ĐNA nhưng hệ thống SKTT khác.'],
    },
    '10_Pham_2024_VN_SocialSupport.docx': {
        'strengths': ['J Affective Disorders Q1. Mixed methods ở VN.', 'Hỗ trợ xã hội = yếu tố bảo vệ — phù hợp Wen 2020 (OR=0,562).'],
        'limits': ['Mẫu SSF (Social Service Facilities) — không đại diện VTN chung.', 'Chỉ Huế.'],
        'gaps': ['NC hỗ trợ xã hội ở VTN trường học VN (không chỉ SSF).'],
    },
    '11_Saikia_2023_IJCM.docx': {
        'strengths': ['IJCM. n = 400, Ấn Độ (Nam Á — gần bối cảnh LMIC).'],
        'limits': ['Chỉ Assam, Ấn Độ. DASS-21 sàng lọc.'],
        'gaps': ['So sánh VN vs Ấn Độ (cùng LMIC) — yếu tố văn hóa khác.'],
    },
    '12_Wen_2020_IJERPH.docx': {
        'strengths': ['IJERPH Q1. LPA — phương pháp nâng cao.', 'Áp lực học tập OR=11,6 — yếu tố mạnh nhất.', 'Hỗ trợ SKTT trường OR=0,562 — bảo vệ.'],
        'limits': ['Chỉ TQ nông thôn. DASS-21 sàng lọc.', 'Cắt ngang — LPA mô tả nhóm, không nhân quả.'],
        'gaps': ['LPA tương tự cho VN — xác định nhóm lo âu cao. Norway 2025 (QT21): bất mãn trường giải thích chính.'],
    },
    '13_Xu_2021_JAffectDisord.docx': {
        'strengths': ['J Affective Disorders Q1. VTN TQ.'],
        'limits': ['Chỉ TQ. Cắt ngang.'],
        'gaps': ['So sánh yếu tố gia đình VN vs TQ — Vĩnh Lộc 2024 (VN20): cha mẹ ly hôn tăng nguy cơ.'],
    },
}

# Process each file
fixed = 0
for filename, crit in critiques.items():
    filepath = os.path.join(OUT, filename)
    if not os.path.exists(filepath):
        print(f'  SKIP: {filename} not found')
        continue

    doc = docx.Document(filepath)
    text = '\n'.join([p.text for p in doc.paragraphs])

    # Check if already has critique
    if 'Phản biện' in text or 'CRITICAL' in text or 'PHAN BIEN' in text:
        print(f'  SKIP: {filename} already has critique')
        continue

    # Add critique at end
    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in crit['strengths']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế chi tiết:', bold=True)
    for s in crit['limits']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Khoảng trống nghiên cứu / Gap:', bold=True)
    for s in crit['gaps']:
        add_red(doc, f'• {s}')

    doc.save(filepath)
    fixed += 1
    d2 = docx.Document(filepath)
    t2 = '\n'.join([p.text for p in d2.paragraphs])
    print(f'  ADDED: {filename} ({len(t2)} chars)')

print(f'\nFixed: {fixed} files')
