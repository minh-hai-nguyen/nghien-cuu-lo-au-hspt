# -*- coding: utf-8 -*-
"""Index các tài liệu MỚI 27/04/2026 vào RAG + KG.
Tài liệu:
- TuLieu_NN_Coping_LoAu_HSTH_v3_chi_tiet_27042026.docx (15 nghiên cứu mới)
- Tra_loi_3_cau_hoi_cho_thay_v2_TiengViet_27042026.docx (3 câu hỏi của thầy)
- BaoCao_KiemTra_3BanDich_Brown_27042026.docx (báo cáo kiểm 10 vòng)
"""
import sys, io, os, json, re
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document

DOC_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao'
KG_DIR  = r'c:\Users\OS\OneDrive\read_books\Lo-au\06_Scripts\kg_data'
RAG_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\rag_dich_phan_bien'
COLLECTION = 'lo_au_dich_phan_bien'

DOCS = [
    {
        'file': 'TuLieu_NN_Coping_LoAu_HSTH_v3_chi_tiet_27042026.docx',
        'doc_id': 'TLNN_Coping_v3',
        'descriptor': 'TuLieu_NuocNgoai_Coping_LoAu_HSTH_v3_27042026',
        'topic': 'foreign_literature_coping_anxiety_secondary_school',
        'type_doc': 'literature_compilation',
        'date': '2026-04-27',
        'language': 'Vietnamese_with_EN_glossary',
        'n_papers_summarized': 15,
    },
    {
        'file': 'Tra_loi_3_cau_hoi_cho_thay_v2_TiengViet_27042026.docx',
        'doc_id': 'TraLoi_3CauHoi_v2',
        'descriptor': 'Tra_loi_3_cau_hoi_cho_thay_v2_TiengViet_27042026',
        'topic': 'three_questions_qualitative_coping_Jefferies_BESST_80percent',
        'type_doc': 'qa_for_thay',
        'date': '2026-04-27',
        'language': 'Vietnamese_pure',
        'questions': 3,
    },
    {
        'file': 'BaoCao_KiemTra_3BanDich_Brown_27042026.docx',
        'doc_id': 'BaoCao_KiemTra_3BanDich',
        'descriptor': 'BaoCao_KiemTra_3BanDich_Brown_10vong_27042026',
        'topic': 'verification_report_3_translations_Brown',
        'type_doc': 'verification_report',
        'date': '2026-04-27',
        'language': 'Vietnamese_pure',
    },
]

# ==== BƯỚC 1 — CHUNKING ====
print('='*60)
print('BƯỚC 1 — CHUNKING các doc mới')
print('='*60)

def docx_to_sections(path):
    d = Document(path)
    chunks = []
    current_h = ''
    current_text = []
    for p in d.paragraphs:
        is_heading = False
        for r in p.runs:
            if r.bold and r.font.size and r.font.size.pt >= 13:
                is_heading = True; break
        text = p.text.strip()
        if not text: continue
        if is_heading and (text.startswith('PHẦN') or text.startswith('NHÓM') or
                           text.startswith('CÂU') or re.match(r'^\d+\.', text) or
                           text.startswith('Bài') or text.isupper()):
            if current_text:
                chunks.append((current_h, '\n'.join(current_text)))
            current_h = text
            current_text = [text]
        else:
            current_text.append(text)
    if current_text:
        chunks.append((current_h, '\n'.join(current_text)))
    for ti, t in enumerate(d.tables):
        rows_txt = []
        for row in t.rows:
            rows_txt.append(' | '.join(c.text.strip() for c in row.cells))
        chunks.append((f'BẢNG {ti+1}', '\n'.join(rows_txt)))
    chunks = [(h, t) for h, t in chunks if len(t) > 200]
    return chunks

all_chunks = []
for doc_meta in DOCS:
    path = os.path.join(DOC_DIR, doc_meta['file'])
    if not os.path.exists(path):
        print(f'  ✗ NOT FOUND: {path}')
        continue
    chunks = docx_to_sections(path)
    print(f'  {doc_meta["doc_id"]}: {len(chunks)} chunks')
    for hi, (h, t) in enumerate(chunks):
        all_chunks.append({
            'chunk_id': f'{doc_meta["doc_id"]}_chunk_{hi:02d}',
            'doc_id': doc_meta['doc_id'],
            'section': h[:80],
            'text': t[:3500],
            'metadata': doc_meta,
        })
print(f'  TỔNG: {len(all_chunks)} chunks')

# ==== BƯỚC 2 — ADD TO RAG ====
print()
print('='*60)
print('BƯỚC 2 — Index vào RAG (rag_dich_phan_bien / lo_au_dich_phan_bien)')
print('='*60)

os.environ['HF_HOME'] = r'D:\hf_cache'
import chromadb
from sentence_transformers import SentenceTransformer
model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')
client = chromadb.PersistentClient(path=RAG_DIR)
col = client.get_collection(COLLECTION)
print(f'  Trước index: {col.count()} chunks')

ids = [c['chunk_id'] for c in all_chunks]
docs_text = []
metas = []
for c in all_chunks:
    full_text = (
        f"DOC {c['doc_id']} — {c['metadata']['descriptor']}\n"
        f"Topic: {c['metadata']['topic']}\n"
        f"Type: {c['metadata']['type_doc']} | Date: {c['metadata']['date']}\n"
        f"--- SECTION: {c['section']} ---\n"
        f"{c['text']}"
    )
    docs_text.append(full_text)
    md = {
        'doc_id': c['doc_id'],
        'section': c['section'],
        'topic': c['metadata']['topic'],
        'type_doc': c['metadata']['type_doc'],
        'date': c['metadata']['date'],
        'language': c['metadata']['language'],
        'source': 'tailieu_moi_27042026',
    }
    metas.append(md)

print(f'  Encoding {len(ids)} chunks...')
emb = model.encode(docs_text, show_progress_bar=False).tolist()

existing = col.get(ids=ids)
if existing['ids']:
    col.delete(ids=existing['ids'])
    print(f'  Deleted {len(existing["ids"])} existing chunks (rerun)')

col.add(ids=ids, embeddings=emb, documents=docs_text, metadatas=metas)
print(f'  Sau index: {col.count()} chunks (added {len(ids)})')

# ==== BƯỚC 3 — UPDATE KG ====
print()
print('='*60)
print('BƯỚC 3 — Update KG nodes + edges')
print('='*60)

with open(os.path.join(KG_DIR, 'nodes.json'), 'r', encoding='utf-8') as f:
    nodes = json.load(f)
with open(os.path.join(KG_DIR, 'edges.json'), 'r', encoding='utf-8') as f:
    edges = json.load(f)
print(f'  Trước: {len(nodes)} nodes, {len(edges)} edges')

new_nodes = []
new_edges = []

# Add doc nodes
for doc_meta in DOCS:
    did = doc_meta['doc_id']
    if any(n.get('id') == did for n in nodes): continue
    new_nodes.append({
        'id': did,
        'type': 'Document',
        'descriptor': doc_meta['descriptor'],
        'file': doc_meta['file'],
        'doc_type': doc_meta['type_doc'],
        'topic': doc_meta['topic'],
        'date': doc_meta['date'],
        'language': doc_meta['language'],
    })
    new_edges.append({'source': did, 'target': f'Date::{doc_meta["date"]}', 'rel': 'CREATED_ON'})
    new_edges.append({'source': did, 'target': f'Topic::{doc_meta["topic"]}', 'rel': 'ABOUT_TOPIC'})

# Add 15 paper nodes (from TuLieu Coping v3) - simplified list
PAPERS_15 = [
    ('TLN_001', 'Frontiers Psychiatry 2025 SR-MA 38 RCTs', 'systematic_review_meta_analysis',
     '10.3389/fpsyt.2025.1594658', 'PMC12127306', 'school_resilience'),
    ('TLN_002', 'Springer Child Psychiatry HumDev 2024 SR 31 studies',
     'systematic_review', '10.1007/s10578-024-01667-5', 'PMC12628395',
     'academic_stress_HS'),
    ('TLN_003', 'Vogelaar 2024 Stress Lessons cluster RCT',
     'cluster_RCT', '', '', 'stress_management_NL'),
    ('TLN_004', 'Lochman 2025 Coping Power early adolescent', 'cluster_RCT',
     '', '', 'coping_power_USA'),
    ('TLN_005', 'MDPI Pediatric Reports 2025 MBI SR-MA 14 studies',
     'systematic_review_meta_analysis', '', '', 'mindfulness_MBI'),
    ('TLN_006', 'PMC12173555 2025 MBI 8+4 weeks booster RCT', 'RCT',
     '', 'PMC12173555', 'mindfulness_booster'),
    ('TLN_007', 'Fulambarkar 2023 cautionary MA mindfulness school',
     'meta_analysis_critical', '10.1111/camh.12572', '',
     'mindfulness_cautionary'),
    ('TLN_008', 'School Mental Health 2023 MBI 10 weeks',
     'controlled_study', '10.1007/s12310-023-09620-y', '', 'mindfulness_10w'),
    ('TLN_009', 'Springer Admin Policy MH 2024 ER MA d=0.37',
     'meta_analysis', '10.1007/s10488-024-01373-3', '', 'emotion_regulation'),
    ('TLN_010', 'Frontiers Psychiatry 2024 ER skills training adol+parents',
     'feasibility', '10.3389/fpsyt.2024.1448529', '', 'ER_parents'),
    ('TLN_011', 'JMIR Pediatrics 2024 Unguided iCBT subthreshold SAD multicenter RCT',
     'multicenter_RCT', '10.2196/55786', '', 'iCBT_SAD'),
    ('TLN_012', 'JMIR Serious Games 2022 Digital interventions ER SR-MA',
     'systematic_review_meta_analysis', '10.2196/31456', '',
     'digital_games'),
    ('TLN_013', 'Murphy 2024 peer support scoping review',
     'scoping_review', '10.1002/jcop.23090', '', 'peer_support'),
    ('TLN_014', 'NCBI Bookshelf NBK602668 UK gov peer support youth',
     'government_review', '', 'NBK602668', 'peer_support_UK'),
    ('TLN_015', 'Lancet RH Western Pacific 2024 China school MH scoping',
     'scoping_review', '', '', 'china_school_MH'),
]

for pid, descriptor, ptype, doi, pmc, topic in PAPERS_15:
    if any(n.get('id') == pid for n in nodes): continue
    node = {
        'id': pid, 'type': 'Paper', 'descriptor': descriptor,
        'paper_type': ptype, 'topic': topic, 'source_doc': 'TLNN_Coping_v3',
    }
    if doi: node['doi'] = doi
    if pmc: node['pmc'] = pmc
    new_nodes.append(node)
    new_edges.append({'source': pid, 'target': 'TLNN_Coping_v3', 'rel': 'CITED_IN_DOC'})
    if doi: new_edges.append({'source': pid, 'target': f'DOI::{doi}', 'rel': 'HAS_DOI'})
    if pmc: new_edges.append({'source': pid, 'target': f'PMC::{pmc}', 'rel': 'HAS_PMC'})
    new_edges.append({'source': pid, 'target': f'Topic::{topic}', 'rel': 'ABOUT_TOPIC'})

# Cross-references: doc TraLoi → BESST + PLACES (đã có trong KG)
new_edges.extend([
    {'source': 'TraLoi_3CauHoi_v2', 'target': 'QT042_BESST', 'rel': 'ANSWERS_ABOUT'},
    {'source': 'TraLoi_3CauHoi_v2', 'target': 'QT_PLACES', 'rel': 'ANSWERS_ABOUT'},
    {'source': 'TraLoi_3CauHoi_v2', 'target': 'Concept::self_referral', 'rel': 'EXPLAINS_CONCEPT'},
    {'source': 'TraLoi_3CauHoi_v2', 'target': 'Concept::Photovoice', 'rel': 'INTRODUCES_CONCEPT'},
    {'source': 'TraLoi_3CauHoi_v2', 'target': 'Concept::bottom_up_research', 'rel': 'INTRODUCES_CONCEPT'},
    {'source': 'BaoCao_KiemTra_3BanDich', 'target': 'QT042_BESST', 'rel': 'VERIFIES'},
    {'source': 'BaoCao_KiemTra_3BanDich', 'target': 'QT_PLACES', 'rel': 'VERIFIES'},
    {'source': 'BaoCao_KiemTra_3BanDich', 'target': 'QT042_B5', 'rel': 'VERIFIES'},
    {'source': 'TLNN_Coping_v3', 'target': 'QT042_BESST', 'rel': 'COMPLEMENTS'},
    {'source': 'TLNN_Coping_v3', 'target': 'QT_PLACES', 'rel': 'COMPLEMENTS'},
    {'source': 'TLNN_Coping_v3', 'target': 'Concept::CBT_school_intervention', 'rel': 'OVERVIEWS'},
    {'source': 'TLNN_Coping_v3', 'target': 'Concept::emotion_regulation', 'rel': 'OVERVIEWS'},
    {'source': 'TLNN_Coping_v3', 'target': 'Concept::mindfulness', 'rel': 'OVERVIEWS'},
    {'source': 'TLNN_Coping_v3', 'target': 'Concept::peer_support', 'rel': 'OVERVIEWS'},
])

nodes.extend(new_nodes)
edges.extend(new_edges)

with open(os.path.join(KG_DIR, 'nodes.json'), 'w', encoding='utf-8') as f:
    json.dump(nodes, f, ensure_ascii=False, indent=2)
with open(os.path.join(KG_DIR, 'edges.json'), 'w', encoding='utf-8') as f:
    json.dump(edges, f, ensure_ascii=False, indent=2)

print(f'  Sau: {len(nodes)} nodes (+{len(new_nodes)}), {len(edges)} edges (+{len(new_edges)})')

# ==== BƯỚC 4 — VERIFICATION ====
print()
print('='*60)
print('BƯỚC 4 — Verification queries')
print('='*60)

# RAG queries
print('--- RAG queries ---')
queries = [
    ('Photovoice học sinh ứng phó tự phát giác quan sáng tạo', 'TraLoi_3CauHoi_v2'),
    ('Jefferies Ungar 2020 Vietnam Edelman Unilever CLEAR market research', 'TraLoi_3CauHoi_v2'),
    ('80% học sinh BESST chưa từng tham vấn GP non-consulter', 'TraLoi_3CauHoi_v2'),
    ('Quản lý căng thẳng dựa trên trường học CBT 38 RCTs Frontiers', 'TLNN_Coping_v3'),
    ('Chánh niệm MBSR SMD −0,14 cảnh báo Fulambarkar MYRIAD', 'TLNN_Coping_v3'),
    ('Đào tạo điều hoà cảm xúc d=0,37 DBT ACT', 'TLNN_Coping_v3'),
    ('CBT online iCBT subthreshold lo âu xã hội thanh niên', 'TLNN_Coping_v3'),
    ('Hỗ trợ ngang hàng peer support Murphy scoping review', 'TLNN_Coping_v3'),
    ('Báo cáo kiểm 10 vòng 3 bản dịch Brown phát hiện lỗi PLACES', 'BaoCao_KiemTra_3BanDich'),
]
correct = 0
for q, expected in queries:
    emb_q = model.encode([q]).tolist()
    r = col.query(query_embeddings=emb_q, n_results=3, where={'doc_id': expected},
                  include=['metadatas','distances'])
    if r['metadatas'] and r['metadatas'][0]:
        top = r['metadatas'][0][0]
        match = '✓' if top['doc_id'] == expected else ' '
        if match == '✓': correct += 1
        sec = top['section'][:50]
        did = top['doc_id']
        qprev = q[:55]
        print(f'  {match} "{qprev}" → {did} | section: {sec}')
print(f'  Top-1 accuracy với metadata filter: {correct}/{len(queries)} ({100*correct/len(queries):.0f}%)')

# KG queries
print()
print('--- KG queries ---')
new_doc_nodes = [n for n in nodes if n.get('id', '').startswith(('TLNN_', 'TraLoi_', 'BaoCao_K'))]
print(f'  Doc nodes mới: {len(new_doc_nodes)}')
for n in new_doc_nodes:
    edges_from = [e for e in edges if e['source']==n['id']]
    edges_to = [e for e in edges if e['target']==n['id']]
    nid = n['id']
    print(f'    {nid}: {len(edges_from)} outgoing + {len(edges_to)} incoming edges')

paper_nodes_new = [n for n in nodes if n.get('id', '').startswith('TLN_')]
print(f'  Paper nodes mới (15 nghiên cứu): {len(paper_nodes_new)}')
print(f'  Sample paper với edges:')
for n in paper_nodes_new[:3]:
    edges_from = [e for e in edges if e['source']==n['id']]
    nid = n['id']
    desc = n['descriptor'][:40]
    print(f'    {nid} ({desc}): {len(edges_from)} edges')

print()
print('='*60)
print('HOÀN TẤT INDEX + VERIFICATION')
print('='*60)
