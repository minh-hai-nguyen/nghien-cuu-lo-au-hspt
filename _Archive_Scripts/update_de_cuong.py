# -*- coding: utf-8 -*-
"""Cập nhật đề cương — thêm bài 16-35 + cập nhật đề cương 3 giai đoạn"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.join(os.path.dirname(os.path.abspath(__file__)), '03_Ban-dich'))
from tao_dich_template import *

doc = create_doc()

# ========== TRANG BÌA ==========
add_heading(doc, 'BÁO CÁO TỔNG HỢP VÀ ĐỀ CƯƠNG NGHIÊN CỨU\nLO ÂU Ở HỌC SINH TRUNG HỌC CƠ SỞ VÀ TRUNG HỌC PHỔ THÔNG', 1)
add_p(doc, '(Việt Nam — Đông Nam Á — Thế giới, 2020–2026)', bold=True)
add_p(doc, 'Cập nhật: 05/04/2026. Phong cách viết: thuật toán Công Thị Hằng v5.')
add_p(doc, 'Dựa trên tổng hợp 35 bài nghiên cứu (7 VN + 28 QT). 35 bản dịch + 37 tóm tắt.', bold=True)
add_p(doc, '')

# ========== PHẦN I: TỔNG HỢP 35 BÀI ==========
add_heading(doc, 'PHẦN I — TỔNG HỢP 35 BÀI NGHIÊN CỨU', 1)

add_heading(doc, 'Bảng tổng hợp 35 bài NC', 2)
add_table(doc,
    ['#', 'Tác giả (năm)', 'Quốc gia', 'n', 'Công cụ', 'Phát hiện chính'],
    [
        ['VN01', 'Hoa et al. 2024', 'VN — Hà Nội', '3.910', 'GAD-7', 'Lo âu 40,6%; nữ > nam'],
        ['VN02', 'V-NAMHS 2022', 'VN — quốc gia', '5.996', 'DISC-5', 'Chẩn đoán 2,3%; tiếp cận 8,4%'],
        ['VN03', 'Pham et al. 2024', 'VN — Huế', '500', 'DASS-21', 'Hỗ trợ XH bảo vệ lo âu'],
        ['VN14', 'Hoàng Trung Học 2025', 'VN', '2.000', 'DASS-21', 'Lo âu 41,5%→25,4% (hồi phục)'],
        ['VN15', 'Ngô Anh Vinh 2024', 'VN — DTTS', '845', 'DASS-21+ACEs', 'Lo âu 54,4%; ACEs liên quan'],
        ['VN16', 'Bảo Quyên 2025', 'VN — Hà Nội', '501', 'DASS-21', 'Lo âu 86,2% — cao nhất'],
        ['VN17', 'Danh Lâm 2022', 'VN — Thanh Hóa', '482', 'DASS-21', 'Lo âu 49%; tự hại 10%'],
        ['VN18', 'An Giang 2025', 'VN — An Giang', '366', 'DASS-21', 'Lo âu 61,2%'],
        ['VN19', 'Thảo Vi 2025', 'VN — Huế', '685', 'DASS-21+LOT-R', 'Lo âu 65,8%; lạc quan trung gian'],
        ['VN20', 'Vĩnh Lộc 2024', 'VN — TPHCM', '976', 'DASS-Y', 'Lo âu 25,1% (DASS-Y thấp hơn)'],
        ['QT01', 'Jenkins 2023', 'Mỹ', '92', 'PHQ-9A+GAD-10', 'Lo âu 50,6%; trầm cảm 44%'],
        ['QT02', 'Saikia 2023', 'Ấn Độ', '400', 'DASS-21', 'Lo âu 30,5%'],
        ['QT03', 'Zhameden 2025', 'LMIC', 'Review', 'Tổng quan HT', 'Can thiệp trường: 1/4 hiệu quả lo âu'],
        ['QT04', 'Anderson 2025', 'Toàn cầu', 'Review', 'Tổng quan', 'Đa yếu tố: MXH, COVID, giáo dục'],
        ['QT05', 'Zhu 2025', 'TQ', '5.000+', 'DASS-21', 'Giấc ngủ <5h: AOR=13,71'],
        ['QT06', 'GBD ASEAN 2025', 'ASEAN', 'GBD', 'GBD 2021', '80,4M ca; VN thứ 3 ASEAN'],
        ['QT07', 'Puyat 2025', 'Philippines', 'Quốc gia', 'GSHS', 'Trầm cảm gấp đôi; tiếp cận 2%'],
        ['QT08', 'Wen 2020', 'TQ nông thôn', '3.000+', 'DASS-21+LPA', 'Áp lực OR=11,6; trường OR=0,562'],
        ['QT09', 'Xu 2021', 'TQ', '2.000+', 'DASS-21', 'Lo âu liên quan gia đình, MXH'],
        ['QT10', 'GBD ASEAN', 'ASEAN', 'GBD', 'GBD 2021', 'VN đứng thứ 3'],
        ['QT21', 'Norway 2025', 'Na Uy', '979K', 'HSCL-6', 'Trường+MXH giải thích xu hướng'],
        ['QT22', 'Li 2025', 'Úc', '4.058', 'PHQ-A+CAS-8', 'Dọc yếu; lo âu p=0,443'],
        ['QT23', 'JAACAP 2024', 'Mỹ', '13,7M', 'Chẩn đoán', 'Lo âu GẤP ĐÔI (AOR=2,17)'],
        ['QT24', 'WHO Europe 2025', 'Châu Âu', '53 nước', 'Policy review', '9M VTN có RLSKTT'],
        ['QT25', 'Crisp 2025', 'Úc', '5.656', 'K10+MHC-SF', 'Flourishing giảm 53%→44,4%'],
        ['QT26', 'UK NHS 2024', 'Anh', 'Quốc gia', 'NHS Digital', 'RLTT 12%→20%; £16 tỷ/năm'],
        ['QT27', 'Fassi 2025', 'UK', '3.340', 'DAWBA', 'Nội hóa nhạy MXH hơn (g=0,46)'],
        ['QT28', 'Zugman 2024', 'Review', '—', 'Tổng quan', 'CBT+SSRI 80,7% (CAMS)'],
        ['QT29', 'Li NMA 2025', '12 nước', '1.711', 'NMA 30 RCT', 'ACT hạng 1; CBT hạng 2'],
        ['QT30', 'Zhang 2025', '204 nước', 'GBD', 'Joinpoint', 'AAPC 0,84%; 10-14 nhanh nhất'],
        ['QT31', 'Islam 2025', '59 nước', '179K', 'GSHS', 'Thực phẩm AOR=2,22; tự tử 2,84'],
        ['QT32', 'Fitzgerald 2024', 'Ireland', '15.334', 'DASS-21', 'Lo âu tăng 48% (2012→2019)'],
        ['QT33', 'JAMA 2024', 'Đan Mạch', '181', 'SDQ (RCT)', 'Giảm ST → d=0,53 nội hóa'],
        ['QT34', 'Cho 2024', 'Hàn Quốc', '1,14M', 'KYRBS', 'Đảo chiều COVID; nghèo 62,8%'],
        ['QT35', 'Jefferies 2020', '7 nước', '6.825', 'SIAS-17', 'VN SAD=30,7%; 18% ẩn'],
    ],
    widths=[1.0, 2.5, 2.0, 1.5, 2.0, 5.0])

# ========== PHẦN II: TỔNG HỢP LIÊN BÀI ==========
add_heading(doc, 'PHẦN II — TỔNG HỢP LIÊN BÀI (xem file riêng: Tổng hợp liên bài báo - Lo âu HS - 04042026.docx)', 1)
add_p(doc, '9 bảng tổng hợp: (1) Tỷ lệ theo quốc gia, (2) Xu hướng 7 NC dài hạn, (3) Nguyên nhân 10 yếu tố, (4) Giới tính 7 NC, (5) Can thiệp 9 phương pháp, (6) Screen time 5 NC, (7) Dịch vụ SKTT 6 nước, (8) Top 10 Gap, (9) Đề cương gợi ý 3 giai đoạn.')

# ========== PHẦN III: ĐỀ CƯƠNG ==========
add_heading(doc, 'PHẦN III — ĐỀ CƯƠNG NGHIÊN CỨU', 1)
add_heading(doc, 'LO ÂU Ở HỌC SINH TRUNG HỌC CƠ SỞ VÀ TRUNG HỌC PHỔ THÔNG TẠI VIỆT NAM', 2)
add_p(doc, 'Đề cương dựa trên tổng hợp 35 bài NC (2020–2025). Cập nhật 05/04/2026.', italic=True)

# 1. Đặt vấn đề
add_heading(doc, '1. Đặt vấn đề và tổng quan tài liệu', 2)
add_p(doc, 'Lo âu là rối loạn tâm thần phổ biến nhất ở thanh thiếu niên toàn cầu: 4,4% trẻ 10–14 tuổi và 5,5% trẻ 15–19 tuổi (WHO, 2024). Xu hướng TĂNG liên tục trên toàn cầu: AAPC 0,84%/năm (GBD 2025, Zhang et al., QT30), lo âu tăng gấp đôi tại Mỹ (JAACAP 2024, Mojtabai & Olfson, QT23: AOR = 2,17), tăng 48% tại Ireland (Fitzgerald et al. 2024, QT32), và tăng liên tục 13 năm tại Na Uy (Brunborg et al. 2025, QT21).')
add_p(doc, 'Tại Việt Nam, tỷ lệ lo âu ở VTN dao động từ 2,3% (chẩn đoán DISC-5, V-NAMHS 2022) đến 86,2% (sàng lọc DASS-21, Bảo Quyên 2025) — khoảng cách gấp 37 lần phản ánh khác biệt CÔNG CỤ ĐO, không phải khác biệt lo âu thực sự. Lo âu xã hội riêng: 30,7% thanh niên VN (Jefferies & Ungar, 2020, QT35 — 1/6 "lo âu ẩn"). Nhóm DTTS bị ảnh hưởng nặng: 54,4% (Ngô Anh Vinh 2024, VN15).')
add_p(doc, 'Các yếu tố nguy cơ đã xác định: bất mãn trường học (Norway 2025, QT21: giải thích chính), áp lực học tập (Wen 2020: OR = 11,6; Vĩnh Lộc 2024, VN20: ESSA ≥59), mạng xã hội (Nature 2025, Fassi et al., QT27: nội hóa g = 0,46), cấu trúc gia đình (VN20: cha mẹ ly hôn), bất an thực phẩm (59 Countries, Islam et al. 2025, QT31: AOR = 2,22), và bất bình đẳng thu nhập (Korea 2024, Cho et al., QT34: stress nghèo 62,8% vs giàu 40,1%).')
add_p(doc, 'Yếu tố bảo vệ: hỗ trợ trường học (Wen 2020: OR = 0,562), cha mẹ tham gia (59 Countries: AOR = 0,75), One Good Adult (Ireland 2024, QT32), lạc quan (Thảo Vi 2025, VN19: β gián tiếp = −0,24).')
add_p(doc, 'Về can thiệp: CBT hiệu quả nhất (AJP 2024, Zugman et al., QT28: CBT+SSRI 80,7% đáp ứng; BMC NMA 2025, Li et al., QT29: CBT SUCRA 0,66). ACT xếp hạng 1 NMA (SUCRA 0,69). Hoạt động thể chất cũng hiệu quả (SUCRA 0,51). Giảm screen time cải thiện SKTT (JAMA 2024, Schmidt-Persson et al., QT33: Cohen d = 0,53). Tuy nhiên, VN CHƯA CÓ BẤT KỲ RCT nào (Zhameden 2025, QT03: 0 RCT từ VN).')
add_p(doc, 'KHOẢNG TRỐNG lớn nhất: (1) 0 RCT can thiệp SKTT tại trường VN, (2) Chưa so sánh sàng lọc vs chẩn đoán trên cùng mẫu, (3) Thiếu dữ liệu xu hướng dài hạn, (4) Chưa phân tích SKTT theo thu nhập, (5) Chưa NC lo âu xã hội riêng ở VTN VN.')

# 2. Mục tiêu
add_heading(doc, '2. Mục tiêu nghiên cứu', 2)
add_p(doc, '2.1. Mục tiêu tổng quát:', bold=True)
add_p(doc, 'Đánh giá thực trạng lo âu, yếu tố liên quan, và hiệu quả can thiệp ở học sinh THCS và THPT tại Việt Nam.')
add_p(doc, '2.2. Mục tiêu cụ thể:', bold=True)
add_p(doc, '(1) Xác định tỷ lệ lo âu (GAD-7, DASS-Y, SIAS-17) và so sánh công cụ đo ở HS THCS/THPT 3 vùng VN.')
add_p(doc, '(2) Phân tích yếu tố nguy cơ (áp lực học tập, MXH, gia đình, thực phẩm, bắt nạt) và bảo vệ (hỗ trợ trường, OGA, lạc quan, hoạt động thể chất).')
add_p(doc, '(3) Đánh giá hiệu quả can thiệp CBT nhóm + hoạt động thể chất tại trường bằng RCT cụm.')
add_p(doc, '(4) Khám phá trải nghiệm lo âu và rào cản tiếp cận dịch vụ qua phỏng vấn sâu.')

# 3. Phương pháp
add_heading(doc, '3. Phương pháp nghiên cứu', 2)
add_p(doc, 'Thiết kế hỗn hợp tuần tự giải thích (sequential explanatory mixed-method), 3 giai đoạn:')

add_heading(doc, '3.1. Giai đoạn 1 — Khảo sát cắt ngang (năm 1)', 3)
add_p(doc, 'Mẫu: ≥1.500 HS THCS + THPT, 3 vùng (đô thị: Hà Nội/TPHCM; nông thôn: Thanh Hóa/An Giang; DTTS: Lạng Sơn/Đắk Lắk). Lấy mẫu cụm nhiều bậc.')
add_p(doc, 'Công cụ:', bold=True)
add_p(doc, '• GAD-7 (so sánh Hoa 2024, VN01: 40,6% Hà Nội)')
add_p(doc, '• DASS-Y (so sánh Vĩnh Lộc 2024, VN20: 25,1% TPHCM — phiên bản VTN)')
add_p(doc, '• SIAS-17 (so sánh Jefferies 2020, QT35: 30,7% VN — lo âu xã hội)')
add_p(doc, '• ESSA (áp lực học tập — Vĩnh Lộc 2024, VN20)')
add_p(doc, '• ACEs 10 mục (Ngô Anh Vinh 2024, VN15)')
add_p(doc, '• Screen time (5 loại: MXH, game, TV, học tập, khác — Hoàng Trung Học 2025, VN14)')
add_p(doc, '• Giấc ngủ (Zhu 2025, QT05: AOR = 13,71 cho <5h)')
add_p(doc, '• LOT-R (lạc quan — Thảo Vi 2025, VN19)')
add_p(doc, '• SDQ (Strengths and Difficulties — JAMA 2024, QT33)')
add_p(doc, '• OGA (One Good Adult — Ireland 2024, QT32)')
add_p(doc, 'Mẫu con (n ≈ 300): DISC-5 chẩn đoán (so sánh V-NAMHS 2022, VN02: 2,3%) → xác định khoảng cách sàng lọc–chẩn đoán.')
add_p(doc, 'Phân tích: Hồi quy logistic đa biến (AOR). LPA (Wen 2020, QT08). Phân tầng: giới, vùng, thu nhập (Korea 2024, QT34). Decomposition (Norway 2025, QT21) nếu có 2 thời điểm.')

add_heading(doc, '3.2. Giai đoạn 2 — RCT cụm can thiệp tại trường (năm 2)', 3)
add_p(doc, 'Thiết kế: RCT cụm (cluster RCT) — 5 trường can thiệp, 5 trường đối chứng. ≥500 HS.')
add_p(doc, 'Can thiệp (12 tuần):', bold=True)
add_p(doc, '• CBT nhóm (8–12 HS/nhóm, 1 lần/tuần, 60 phút) — BMC NMA 2025 (QT29): CBT SUCRA 0,66; AJP 2024 (QT28): CBT đơn 59,7% đáp ứng.')
add_p(doc, '• Hoạt động thể chất (PE) tại trường (2 lần/tuần, 45 phút: thể dục + yoga/mindfulness) — BMC NMA 2025 (QT29): PE SUCRA 0,51; cơ chế: endorphin, BDNF, HPA.')
add_p(doc, '• Giảm screen time giải trí toàn gia đình (khuyến khích, không bắt buộc) — JAMA 2024 (QT33): Cohen d = 0,53 chỉ 2 tuần.')
add_p(doc, '• Chương trình OGA (mentor tại trường) — Ireland 2024 (QT32): OGA bảo vệ nhất quán.')
add_p(doc, 'Đo: Trước–sau–3 tháng (GAD-7 + DASS-Y + SDQ + SIAS-17). Phân tích: ITT, mixed-effects models, Cohen d.')
add_p(doc, 'So sánh kỳ vọng:', bold=True)

add_table(doc,
    ['Can thiệp', 'Bằng chứng', 'Kỳ vọng', 'Nguồn'],
    [['CBT nhóm', 'CBT đơn 59,7% đáp ứng', 'Giảm GAD-7 ≥3 điểm', 'AJP 2024 (QT28)'],
     ['PE tại trường', 'PE SUCRA 0,51', 'Giảm DASS-Y ≥2 điểm', 'BMC NMA 2025 (QT29)'],
     ['Giảm screen time', 'Cohen d = 0,53', 'Giảm SDQ nội hóa ≥1 điểm', 'JAMA 2024 (QT33)'],
     ['OGA/Mentor', 'OGA bảo vệ (Ireland)', 'Giảm cả lo âu + trầm cảm', 'Ireland 2024 (QT32)'],
     ['Kết hợp tất cả', 'CBT+SSRI 80,7%', 'Cải thiện ≥50% mẫu can thiệp', 'CAMS (Walkup 2008)']],
    widths=[2.5, 3.0, 3.0, 3.0])

add_heading(doc, '3.3. Giai đoạn 3 — Phỏng vấn sâu (năm 2–3)', 3)
add_p(doc, 'Mẫu: 30–40 HS (chọn có mục đích: lo âu cao/thấp, nam/nữ, đô thị/nông thôn/DTTS). Phỏng vấn bán cấu trúc về: trải nghiệm lo âu, nguồn gốc, rào cản tiếp cận dịch vụ, phản hồi can thiệp.')
add_p(doc, 'Phân tích chủ đề (thematic analysis). Bổ sung cho kết quả định lượng — đặc biệt hiểu "lo âu ẩn" (Jefferies 2020, QT35: 18% có SAD nhưng không biết) và "false negatives" phụ huynh (V-NAMHS 2022: chỉ 5,1% phụ huynh nhận ra).')

# 4. Ý nghĩa
add_heading(doc, '4. Ý nghĩa khoa học và thực tiễn', 2)
add_p(doc, 'Khoa học:', bold=True)
add_p(doc, '• RCT ĐẦU TIÊN can thiệp SKTT tại trường VN — lấp đầy Gap #1 cross-study.')
add_p(doc, '• So sánh GAD-7 vs DASS-Y vs SIAS-17 vs DISC-5 — xác định công cụ phù hợp VN (Gap #2).')
add_p(doc, '• Dữ liệu xu hướng: kết hợp với V-NAMHS 2022 tạo thời điểm 2 (Gap #3).')
add_p(doc, '• Phân tầng thu nhập/vùng — so sánh với Korea 2024 (Gap #4).')
add_p(doc, '• NC lo âu xã hội (SIAS-17) ở VTN VN — so sánh trực tiếp Jefferies 2020 (Gap #5).')
add_p(doc, 'Thực tiễn:', bold=True)
add_p(doc, '• Mô hình can thiệp tại trường khả thi (CBT nhóm + PE + OGA) — có thể nhân rộng.')
add_p(doc, '• Cơ sở dữ liệu cho chính sách SKTT VTN quốc gia.')
add_p(doc, '• Đào tạo giáo viên + cố vấn tâm lý học đường.')

# 5. Tiến độ
add_heading(doc, '5. Tiến độ dự kiến', 2)
add_table(doc,
    ['Giai đoạn', 'Thời gian', 'Nội dung', 'Sản phẩm'],
    [['Chuẩn bị', 'Tháng 1–3 năm 1', 'Dịch/xác thực công cụ; phê duyệt đạo đức', 'Bộ công cụ tiếng Việt'],
     ['GĐ1: Khảo sát', 'Tháng 4–9 năm 1', 'Thu thập + phân tích dữ liệu cắt ngang', '2 bài báo Q1/Q2'],
     ['GĐ2: RCT', 'Tháng 1–12 năm 2', 'Can thiệp 12 tuần + theo dõi 3 tháng', '1 bài báo Q1 (RCT)'],
     ['GĐ3: Phỏng vấn', 'Tháng 6–12 năm 2', 'Phỏng vấn sâu + phân tích chủ đề', '1 bài báo mixed-method'],
     ['Tổng hợp', 'Năm 3', 'Viết luận văn/báo cáo cuối', 'Luận văn + policy brief']],
    widths=[2.0, 2.5, 4.5, 3.0])

# 6. Tài liệu tham khảo chính
add_heading(doc, '6. Tài liệu tham khảo chính (trích)', 2)
refs = [
    'Brunborg, G.S., et al. (2025). Possible explanations for the upward trend in mental distress among adolescents in Norway. Social Science & Medicine, 384, 118528.',
    'Cho, J., et al. (2024). National trends in adolescents\' mental health by income level in South Korea. Scientific Reports.',
    'Fassi, L., et al. (2025). Social media use in adolescents with and without mental health conditions. Nature Human Behaviour, 9, 1283–1299.',
    'Fitzgerald, A., et al. (2024). Exploring changing trends in depression and anxiety among adolescents from 2012 to 2019. Early Intervention in Psychiatry.',
    'Islam, M.A., et al. (2025). Prevalence and factors associated with anxiety among school going adolescents: 59 countries. J Affective Disorders, 393, 120315.',
    'Jefferies, P. & Ungar, M. (2020). Social anxiety in young people: A prevalence study in seven countries. PLOS ONE, 15(9), e0239133.',
    'Li, L.H., et al. (2025). Effects of different interventions on anxiety disorders in children: NMA. BMC Psychiatry, 25, 809.',
    'Li, S.H., et al. (2025). Cross-sectional and longitudinal associations of screen time with adolescent depression and anxiety. BJCP, 64, 873–887.',
    'Mojtabai, R. & Olfson, M. (2024). Trends in mental disorders in children receiving treatment. JAACAP, 64(8).',
    'Schmidt-Persson, J., et al. (2024). Screen media use and mental health: RCT. JAMA Network Open, 7(7), e2419881.',
    'V-NAMHS (2022). Khảo sát Quốc gia về Sức khỏe Tâm thần VTN VN.',
    'Zugman, A., et al. (2024). Current and future approaches to pediatric anxiety disorder treatment. AJP, 181, 189–200.',
    'Zhang, D., et al. (2025). Trends in depressive and anxiety disorders among 10–24 year olds: GBD analysis. J Affective Disorders.',
    '(Xem đầy đủ danh mục 35 bài NC trong Bảng tổng hợp Phần I)',
]
for ref in refs:
    add_p(doc, ref, size=10)

# SAVE
outpath = 'Lo au - Bao cao CTH + De cuong - 05042026.docx'
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
