# -*- coding: utf-8 -*-
"""Tim cac doan trong doc CO HINH HOC giong heading (bold/co chu lon/ngan)
nhung style la Normal -> bi an khoi TOC.
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

candidates = []  # (idx, text, style, size, bold, in_table_or_after_toc)

# Find where TOC table ends (first table)
toc_table = d.tables[0] if d.tables else None
# Position of TOC table doesn't help via python-docx easily.
# Instead: only consider paragraphs AFTER paragraph 100 (post-front-matter)

for i, p in enumerate(d.paragraphs):
    if i < 100:  # Skip front matter
        continue
    txt = p.text.strip()
    if not txt:
        continue
    if len(txt) > 250:
        continue  # Too long for heading
    style = p.style.name if p.style else '?'
    if 'Heading' in style or 'Tiêu đề' in style:
        continue  # Already heading - will be in TOC
    # Check if bold + numbering pattern + short
    bold_score = 0
    italic_score = 0
    sizes = []
    has_runs = False
    for r in p.runs:
        if r.text.strip():
            has_runs = True
            if r.bold: bold_score += 1
            if r.italic: italic_score += 1
            if r.font.size:
                sizes.append(r.font.size.pt)
    if not has_runs:
        continue
    # Check numbering pattern OR caps OR special structural keyword
    has_numbering = bool(re.match(r'^\d+\.\d', txt))
    is_caps = txt.upper() == txt and len(txt.split()) <= 8
    is_structural = any(kw in txt.upper() for kw in
                        ['TIỂU KẾT', 'KẾT LUẬN', 'KIẾN NGHỊ', 'MỞ ĐẦU',
                         'GIẢ THUYẾT 1', 'GIẢ THUYẾT 2', 'GIẢ THUYẾT 3'])
    score = 0
    if has_numbering: score += 2
    if is_caps: score += 2
    if is_structural: score += 3
    if bold_score >= max(1, len(p.runs) - 2): score += 1  # mostly bold
    if sizes and max(sizes) >= 14: score += 1
    if len(txt) < 100: score += 1

    if score >= 2:
        candidates.append({
            'idx': i,
            'text': txt[:120],
            'style': style,
            'bold': bold_score >= max(1, len(p.runs) - 2),
            'sizes': sizes,
            'score': score,
            'has_numbering': has_numbering,
            'is_caps': is_caps,
            'is_structural': is_structural,
        })

print(f"Tim thay {len(candidates)} paragraphs co dau hieu la heading nhung style != Heading")
print()
print(f"{'idx':>5}  {'score':>5}  {'style':>16}  text")
print("-"*100)
for c in candidates:
    print(f"{c['idx']:>5}  {c['score']:>5}  {c['style']:>16}  {c['text']}")
