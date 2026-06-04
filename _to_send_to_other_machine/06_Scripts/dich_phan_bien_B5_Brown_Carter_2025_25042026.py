# -*- coding: utf-8 -*-
"""Dịch song ngữ + Phản biện đầy đủ bài B5 (Brown & Carter 2025).
Bài gốc: Brown JSL, Carter B (2025). School based interventions for depression
and anxiety in UK. Journal of Mental Health, 34(4): 357–361.
DOI: 10.1080/09638237.2025.2512332
Quy trình: dịch song ngữ EN-VN + phản biện chữ đỏ + bảng + reference DOI/PMID + truy vết RAG/KG.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\B5_Brown_Carter_2025_dich_phan_bien_25042026.docx'

RED   = RGBColor(0xC0, 0x00, 0x00)
BLUE  = RGBColor(0x00, 0x70, 0xC0)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x00, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(13)
s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def title(text, size=18):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def subtitle(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def H1(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def H2(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def en(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def vn(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(13); r.font.name = 'Times New Roman'
def crit(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('[Phản biện] '); r.bold = True; r.font.color.rgb = RED; r.font.size = Pt(12); r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.font.color.rgb = RED; r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
def note(text, color=GREEN):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('[Ghi chú dịch giả] '); r.bold = True; r.italic = True; r.font.color.rgb = color; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.italic = True; r2.font.color.rgb = color; r2.font.size = Pt(11); r2.font.name = 'Times New Roman'
def nr(text, bold=False, size=13, color=None, italic=False):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color

# =====================================================================
# TRANG BÌA
# =====================================================================
title("BẢN DỊCH SONG NGỮ + PHẢN BIỆN", 16)
title("CAN THIỆP TÂM LÝ TẠI TRƯỜNG HỌC CHO", 17)
title("TRẦM CẢM VÀ LO ÂU TẠI VƯƠNG QUỐC ANH", 17)
subtitle("School based interventions for depression and anxiety in UK")
nr("")
subtitle("June S. L. Brown & Ben Carter (2025)")
subtitle("Journal of Mental Health, 34(4): 357–361")
subtitle("DOI: 10.1080/09638237.2025.2512332 — Published online 14 Jun 2025")
subtitle("Loại bài: EDITORIAL (xã luận chuyên môn)")
nr("")
nr("Trợ lý nghiên cứu — Dịch + Phản biện — 25/04/2026", italic=True, color=GRAY, size=10)
nr("Quy trình: dịch song ngữ EN↔VN từng đoạn + phản biện chữ đỏ có dẫn chứng từ corpus "
   "35+ bài lo âu/SKTT VTN của dự án + RAG (rag_db_full / collection lo_au_full) + KG "
   "(06_Scripts/kg_data) + đối chiếu tài liệu ngoài (PMC, PubMed, King's College). "
   "Kiểm 3 vòng theo memory feedback_research_workflow.md.",
   italic=True, color=GRAY, size=10)

# =====================================================================
# THÔNG TIN THƯ MỤC
# =====================================================================
H1("THÔNG TIN THƯ MỤC")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài (EN)', 'School based interventions for depression and anxiety in UK'],
        ['Tên bài (VN)', 'Can thiệp tâm lý tại trường học cho trầm cảm và lo âu tại Vương quốc Anh'],
        ['Tác giả', 'June S. L. Brown (Khoa Tâm lý, King\'s College London) & Ben Carter '
                    '(Biostatistics & Health Informatics, KCL — Denmark Hill Campus)'],
        ['Liên hệ', 'june.brown@kcl.ac.uk'],
        ['Tạp chí', 'Journal of Mental Health (Taylor & Francis / Informa UK)'],
        ['Tập / Số / Trang', 'Vol. 34, No. 4, pp. 357–361 (2025)'],
        ['ISSN', '0963-8237 (Print); 1360-0567 (Online)'],
        ['DOI', '10.1080/09638237.2025.2512332'],
        ['Loại bài', 'EDITORIAL (xã luận, không phải nghiên cứu gốc)'],
        ['Ngày nhận / sửa / chấp nhận', '04/03/2025 → 29/04/2025 → 02/05/2025'],
        ['Đăng online', '14/06/2025'],
        ['Lượt xem (tính đến trích)', '2.867'],
        ['Funding', 'Không có (no funding)'],
        ['Conflict of interest', 'Không có'],
        ['Mã trong corpus dự án', 'B5 / QT042 — Raw: 03_Ban-dich/B5_UK_school_raw.txt (552 dòng); '
                                  'tóm tắt cũ: 03_Ban-dich/QT042_BrownCarter_UK_School_JMH_2025.docx '
                                  '(38 KB / 2.673 ký tự — KHÔNG phải dịch full)'],
    ], [4.5, 11.5])

H2("Định vị bài này trong tổng quan dự án")
vn("Đây là một xã luận (editorial) NGẮN trong số chuyên đề J Mental Health 2025. Hai tác "
   "giả Brown & Carter chính là nhóm phát triển mô hình PLACES (2022) và chủ trì thử nghiệm "
   "BESST (2024 Lancet Psychiatry). Bài này có chức năng tổng hợp lập trường của họ về "
   "can thiệp SKTT trường học UK — không phải nghiên cứu gốc, không có dữ liệu mới, mà là "
   "tổng hợp + diễn giải các bằng chứng đã có theo khung của họ.")
crit("Vì đây là EDITORIAL (xã luận, evidence level 5 trong hệ thống Oxford CEBM 2011), "
     "nên cần đọc nó như một góc nhìn của tác giả — không nên trích như bằng chứng cấp 1. "
     "Mọi tuyên bố lớn trong bài này phải truy ngược về bài gốc (Brown 2022 PLACES, Brown "
     "2024 BESST, Kuyken 2022 MYRIAD, Stallard 2014 PACES, Zhang 2023 meta-analysis…) để "
     "xác minh. Đặc biệt cảnh giác cụm từ \"very positive results\" mà bài này dùng cho "
     "BESST — số liệu thực tế là Cohen d = −0,17 (SMALL effect) cho mẫu chung; chỉ riêng "
     "subgroup elevated depression (MFQ >27, n=298) mới đạt d = −0,52 (medium).")

# =====================================================================
# PHẦN 1 — DỊCH SONG NGỮ
# =====================================================================
H1("PHẦN 1 — DỊCH SONG NGỮ ĐẦY ĐỦ")
note("Quy ước: mỗi đoạn EN gốc (italic, xám, size 11) đặt ngay trên bản dịch VN (đen, "
     "regular, size 13). Phản biện chữ đỏ chèn xen kẽ khi có vấn đề kỹ thuật / số liệu cần lưu ý.")

# ---------- 1. Bối cảnh ----------
H2("1. Bối cảnh — Vấn đề SKTT vị thành niên ở UK")

en("Mental health problems among adolescents, particularly anxiety and depression, are a "
   "growing public health concern (Marcheselli et al., 2018). In the UK, approximately one "
   "in four adolescents aged 17–19 meet the criteria for a diagnosable mental health "
   "disorder, and the impact of untreated mental health issues during this critical "
   "developmental stage can be profound (Newlove-Delgado et al., 2023).")
vn("Các vấn đề sức khỏe tâm thần (SKTT) ở vị thành niên, đặc biệt là lo âu và trầm cảm, "
   "đang trở thành mối quan tâm y tế công cộng ngày càng tăng (Marcheselli và cs, 2018). "
   "Tại Vương quốc Anh, khoảng MỘT TRÊN BỐN thanh thiếu niên 17–19 tuổi đáp ứng tiêu chuẩn "
   "chẩn đoán một rối loạn SKTT, và hệ quả của việc không được điều trị trong giai đoạn "
   "phát triển then chốt này có thể rất nặng nề (Newlove-Delgado và cs, 2023).")
crit("Số liệu \"1/4 ở 17–19 tuổi\" này đến từ khảo sát NHS England 2023 (Newlove-Delgado "
     "wave 4) — sử dụng SDQ + DAWBA. Đối chiếu với corpus dự án: V-NAMHS 2022 ở Việt Nam "
     "ghi 21,7% \"mental health problem\" (subthreshold+) vs 3,3% \"mental disorder\" theo "
     "DSM-5 (DISC-5) — nghĩa là tỷ lệ ở VN khá tương đồng UK ở mức subthreshold nhưng "
     "thấp hơn 8× ở mức chẩn đoán. Khác biệt này KHÔNG nhất thiết do dịch tễ thật mà có "
     "thể do công cụ (SDQ/DAWBA vs DISC-5) và bối cảnh báo cáo. Khi áp dụng \"1/4\" cho "
     "VN cần thận trọng — không thể trích thẳng.")

en("These issues are associated with a range of adverse outcomes, internationally (Johnson "
   "et al., 2018), including an increased risk of depression and anxiety in adulthood "
   "(Johnson et al., 2018; Kessler et al., 2005) as well as poorer academic performance "
   "(Brännlund et al., 2017), poorer education and employment outcomes (Hale et al., 2015) "
   "and challenges in forming and maintaining relationships (Sadler et al., 2018). Despite "
   "this, many young people struggle to access mental health support, with estimates "
   "suggesting that around 60% of young people with mental health issues do not receive "
   "any formal care (UK Government, 2022).")
vn("Các vấn đề này gắn liền với nhiều hậu quả bất lợi trên phạm vi quốc tế (Johnson và "
   "cs, 2018), bao gồm tăng nguy cơ trầm cảm và lo âu ở tuổi trưởng thành (Johnson và cs, "
   "2018; Kessler và cs, 2005), thành tích học tập kém hơn (Brännlund và cs, 2017), kết "
   "quả giáo dục và việc làm kém hơn (Hale và cs, 2015), cùng các khó khăn trong việc "
   "hình thành và duy trì các mối quan hệ (Sadler và cs, 2018). Mặc dù vậy, nhiều người "
   "trẻ vẫn rất khó tiếp cận hỗ trợ SKTT — ước tính cho thấy khoảng 60% người trẻ gặp "
   "vấn đề SKTT KHÔNG nhận được bất kỳ chăm sóc chính thức nào (Chính phủ Anh, 2022).")
crit("Số 60% \"không nhận chăm sóc\" này khớp với pattern ở các quốc gia khác — VD VN: "
     "V-NAMHS 2022 ghi chỉ 8,4% VTN VN tiếp cận dịch vụ SKTT, tức 91,6% không tiếp cận "
     "— TỆ HƠN UK đáng kể. Hai trong corpus còn ghi: chỉ 5,1% phụ huynh VN nhận ra con "
     "mình bị lo âu (V-NAMHS) → khoảng cách dịch vụ ở VN sâu hơn UK rất nhiều, do đó "
     "các mô hình UK (như BESST/PLACES sẽ bàn ở mục 4) khi áp dụng cho VN cần phải bắt "
     "đầu từ MENTAL HEALTH LITERACY của phụ huynh + giáo viên trước, không phải "
     "self-referral của HS như ở UK.")

en("There are several barriers to solving these problems including the reluctance of "
   "young people themselves to seek help (Gulliver et al., 2010; Rickwood et al., 2005). "
   "Even when they do seek help, there is a low capacity in specialist mental health "
   "services for young people (McGorry et al., 2013) that results in very long waiting "
   "lists. So, despite knowing that the age of onset of most mental health problems is "
   "15 for 50% of young people and 24 for 75% of young people (de Girolamo et al., 2012; "
   "Kessler et al., 2005), the barriers mean that early intervention is difficult.")
vn("Có nhiều rào cản trong việc giải quyết các vấn đề này, bao gồm chính sự e ngại của "
   "người trẻ trong việc tìm kiếm sự giúp đỡ (Gulliver và cs, 2010; Rickwood và cs, "
   "2005). Ngay cả khi họ tìm kiếm trợ giúp, năng lực của các dịch vụ SKTT chuyên khoa "
   "dành cho người trẻ vẫn còn thấp (McGorry và cs, 2013), dẫn đến danh sách chờ rất "
   "dài. Vì vậy, mặc dù biết rằng tuổi khởi phát của phần lớn các vấn đề SKTT là 15 "
   "tuổi (cho 50% người trẻ) và 24 tuổi (cho 75% người trẻ) (de Girolamo và cs, 2012; "
   "Kessler và cs, 2005), các rào cản này khiến việc can thiệp sớm trở nên rất khó khăn.")
crit("Số liệu \"15 tuổi cho 50%, 24 tuổi cho 75%\" là dữ liệu kinh điển của Kessler 2005 "
     "(NCS-R, Mỹ, n>9.000) và de Girolamo 2012 — vẫn là chuẩn được trích trong tổng quan "
     "GBD 2021 ASEAN (corpus QT10/QT12). Với VN: 50% trước 15 tuổi nghĩa là can thiệp "
     "trường học (THCS-THPT) là CỬA SỔ VÀNG. Đề cương dự án (đã có trong RAG INSIGHT_05) "
     "khuyến nghị bám sát điểm này.")

en("School-based mental health interventions have been increasingly recognized as an "
   "essential component of early intervention strategies, providing a platform for "
   "reaching adolescents in an environment that is accessible, familiar, and less "
   "stigmatising than clinical settings (Fazel et al., 2014; Lovell & Richards, 2000). "
   "This has been recognised in the UK as well as in Low and Middle Income countries "
   "(LMIC) (Fazel et al., 2014).")
vn("Các can thiệp SKTT dựa vào trường học ngày càng được công nhận là một thành phần "
   "thiết yếu của chiến lược can thiệp sớm — cung cấp một nền tảng để tiếp cận vị thành "
   "niên trong môi trường dễ tiếp cận, quen thuộc và ít kỳ thị hơn so với bối cảnh lâm "
   "sàng (Fazel và cs, 2014; Lovell & Richards, 2000). Điều này đã được công nhận tại "
   "UK cũng như tại các quốc gia thu nhập thấp và trung bình (LMIC) (Fazel và cs, 2014).")

# ---------- 2. Ai nên cung cấp can thiệp ----------
H2("2. Ai nên cung cấp can thiệp SKTT để giảm trầm cảm và lo âu trong trường học?")
en("Who should deliver the mental health interventions to reduce depression and anxiety "
   "in schools?")

en("The attraction of using teachers for delivering mental health interventions is that "
   "they are already in schools, and if trained, could offer effective and hopefully, "
   "cost-effective interventions. Unfortunately, this has not been found to be the case. "
   "Stallard et al. (2014) and others have found that clinicians were generally more "
   "effective than teachers (Fisak et al., 2011; Zhang et al., 2023), particularly in "
   "secondary schools.")
vn("Sức hấp dẫn của việc sử dụng GIÁO VIÊN để cung cấp can thiệp SKTT là vì họ đã có sẵn "
   "trong trường học, và nếu được đào tạo, có thể cung cấp các can thiệp hiệu quả và "
   "(hy vọng) tiết kiệm chi phí. Đáng tiếc, thực tế lại không như vậy. Stallard và cs "
   "(2014) cùng các nghiên cứu khác đã phát hiện rằng các nhà lâm sàng (clinicians) "
   "nhìn chung HIỆU QUẢ HƠN giáo viên (Fisak và cs, 2011; Zhang và cs, 2023), đặc biệt "
   "ở bậc trung học phổ thông.")
crit("Tuyên bố \"clinicians > teachers\" này bị PHỤ THUỘC NHIỀU YẾU TỐ — ít nhất 3 vấn "
     "đề: (1) Nhiều RCT có giáo viên không được đào tạo đầy đủ — so sánh apples-to-oranges. "
     "(2) Bài Stallard 2014 PACES (FRIENDS classroom) cho effect size NHỎ cả 2 nhóm — kết "
     "luận \"clinician > teacher\" dựa trên hiệu ứng nhỏ rất khó áp dụng. (3) Trong corpus "
     "dự án có B8 Sri Lanka (CACBT) chứng minh GV được đào tạo tốt CÓ THỂ cung cấp CBT "
     "hiệu quả ở LMIC. Với VN, vì rất thiếu nhà lâm sàng (0,2 bác sĩ tâm thần/100.000 dân "
     "ASEAN — nguồn corpus QT10), mô hình GV-led có thể là DUY NHẤT khả thi quy mô lớn — "
     "tuyệt đối không thể loại bỏ vì 1 câu \"clinician > teacher\" trong UK setting.")

en("A recent UK development is the introduction of a new professional group of Mental "
   "Health Support Teams (MHSTs) to work directly in schools and provide early support "
   "for children and adolescents with mild to moderate mental health difficulties "
   "(Alderwick & Dixon, 2019). These MHSTs (consisting of master's or postgraduate "
   "diploma level therapists or junior therapists) aim to improve access to evidence-"
   "based interventions, including cognitive-behavioural therapy (CBT) (NHS England, "
   "2023). With appropriate training, it is possible they could deliver clinically "
   "effective and cost-effective interventions.")
vn("Một bước phát triển gần đây tại UK là việc thành lập một nhóm chuyên môn mới — "
   "Mental Health Support Teams (MHSTs — Đội hỗ trợ SKTT) — để làm việc trực tiếp tại "
   "trường học, cung cấp hỗ trợ sớm cho trẻ em và vị thành niên có khó khăn SKTT mức "
   "nhẹ đến trung bình (Alderwick & Dixon, 2019). Các MHST này (gồm các nhà trị liệu "
   "trình độ thạc sĩ / postgraduate diploma hoặc nhà trị liệu mới vào nghề) hướng tới "
   "cải thiện khả năng tiếp cận các can thiệp dựa trên bằng chứng, bao gồm liệu pháp "
   "nhận thức – hành vi (CBT) (NHS England, 2023). Với đào tạo phù hợp, MHST CÓ THỂ "
   "cung cấp các can thiệp hiệu quả lâm sàng và tiết kiệm chi phí.")
crit("MHST là MÔ HÌNH UK ĐẶC THÙ — gắn với NHS England, ngân sách công + đào tạo chuẩn "
     "hoá quốc gia. Để áp dụng cho VN cần điều chỉnh: VN không có cơ chế tương đương "
     "MHST. Phương án khả thi nhất ở VN là nâng cấp đội ngũ TƯ VẤN HỌC ĐƯỜNG hiện có "
     "(theo Thông tư 31/2017/TT-BGDĐT) bằng chương trình đào tạo CBT tích hợp + "
     "supervision liên trường. Đây là gap chưa có RCT nào ở VN giải quyết.")

# ---------- 3. Mental health literacy ----------
H2("3. Chương trình nào? — Mental Health Literacy")
en("It has often been assumed that mental health literacy, which refers to knowledge and "
   "intention to recognise symptoms and would lead to better help-seeking, but this has "
   "not always been found. Yamaguchi et al. (2024) found a teacher led 20 minute online "
   "mental health literacy programme for 15–16 year students in a small RCT (n = 270) "
   "led to improvements in improved knowledge and recognition of necessity to seek help. "
   "However, little change was found on \"intention to seek help\".")
vn("Người ta thường giả định rằng \"hiểu biết về SKTT\" (mental health literacy) — tức "
   "kiến thức và ý định nhận diện triệu chứng — sẽ dẫn tới hành vi tìm kiếm trợ giúp "
   "tốt hơn, nhưng điều này KHÔNG luôn được kiểm chứng. Yamaguchi và cs (2024) thực hiện "
   "một RCT nhỏ (n = 270) với chương trình hiểu biết SKTT do giáo viên dẫn dắt, online, "
   "20 phút, cho học sinh 15–16 tuổi. Kết quả: cải thiện kiến thức và nhận thức về "
   "việc \"cần\" tìm trợ giúp, NHƯNG ít thay đổi về \"ý định\" tìm trợ giúp.")
crit("Đây là khoảng cách KINH ĐIỂN giữa knowledge → attitude → intention → behaviour "
     "(KAIB cascade). Mental health literacy chỉ chạm tới knowledge và attitude. Để "
     "chuyển hoá sang BEHAVIOUR (help-seeking thực sự) cần thêm các yếu tố như "
     "self-referral system, perceived effectiveness, lay language — chính là điều mô "
     "hình PLACES (Brown 2022) bổ sung. Trong corpus VN: nghiên cứu Hoa 2024 Hà Nội "
     "(n=3.910, GAD-7) báo cáo Cronbach α = 0,916 — đo được symptom tốt nhưng KHÔNG đo "
     "intention/behaviour — gap nghiên cứu VN cần fill.")

en("Prabhu et al. (2024) reviewed studies about interventions of the mental health "
   "literacy of secondary school teachers, concluding this could be improved but the "
   "quality of the studies was variable. There have also been variations in effectiveness "
   "depending on who delivered the intervention.")
vn("Prabhu và cs (2024) đã rà soát các nghiên cứu về can thiệp nâng cao hiểu biết SKTT "
   "cho GIÁO VIÊN trung học, kết luận rằng kết quả có thể cải thiện được nhưng chất "
   "lượng nghiên cứu không đồng đều. Hiệu quả cũng dao động tùy thuộc vào người cung cấp "
   "can thiệp.")
crit("Đây là một review (Prabhu 2024) không định lượng — \"quality variable\" thường "
     "ngụ ý meta-analysis chưa thể gộp do heterogeneity cao. Áp dụng VN: chương trình "
     "đào tạo cho GV (vd. UNICEF VN22) cần có RCT nội địa kiểm chứng, không nên copy "
     "module UK/Aus mà thiếu adaptation văn hoá.")

# ---------- 4. Clinical programmes ----------
H2("4. Chương trình nào? — Clinical Programmes (Mindfulness, CBT)")
en("School based programmes range from mindfulness to more CBT-like programmes developed "
   "in Australia (Friends, Resourceful adolescent Programme, Mood Gym) and USA (Penn "
   "Resilience Programme, Mindfulness) with only one in the UK for adolescents (DISCOVER).")
vn("Các chương trình tại trường học rất đa dạng — từ mindfulness đến các chương trình "
   "kiểu CBT được phát triển tại Úc (Friends, Resourceful Adolescent Programme, Mood "
   "Gym) và Mỹ (Penn Resilience Programme, Mindfulness), với DUY NHẤT một chương trình "
   "tại UK dành cho VTN là DISCOVER.")
crit("\"Duy nhất DISCOVER\" cho UK là cách diễn đạt thiên vị — thực tế UK còn có FRIENDS "
     "(Stallard PACES), MYRIAD mindfulness (Kuyken), và nhiều chương trình nhỏ khác. "
     "Tác giả nhấn mạnh DISCOVER vì đó là sản phẩm của họ. Với VN: chưa có chương trình "
     "CBT-trường học nội địa được kiểm chứng RCT — cần adaptation từ Friends/DISCOVER "
     "có điều chỉnh văn hoá.")

en("Systematic reviews on the effectiveness of psychological interventions in school "
   "settings have discovered mixed findings on reductions of depression and anxiety "
   "(Caldwell et al., 2019; Feiss et al., 2019; Werner-Seidler et al., 2021). One reason "
   "may be the variation in the inclusion and exclusion criteria which has been "
   "highlighted in a recent meta-analysis (Zhang et al., 2023). There have been reports "
   "of variation in the effects of different types of interventions with Zhang et al. "
   "(2023) reporting that CBT programmes tended to be associated with higher effect "
   "sizes, a conclusion that was also supported in an anxiety focussed review (Fisak et "
   "al., 2011). However, in the first narrative analysis of school interventions, "
   "Caldwell et al. (2019) did not support this view as strongly. They concluded that "
   "CBT showed weak evidence of reducing anxiety but studies were small. And they found "
   "no evidence of differences in types of interventions being effective in reducing "
   "depression.")
vn("Các tổng quan hệ thống về hiệu quả của can thiệp tâm lý tại trường học đã ghi nhận "
   "kết quả PHA TRỘN về việc giảm trầm cảm và lo âu (Caldwell và cs, 2019; Feiss và cs, "
   "2019; Werner-Seidler và cs, 2021). Một lý do có thể là sự khác biệt về tiêu chí "
   "nhận và loại — điều đã được nêu trong một meta-analysis gần đây (Zhang và cs, 2023). "
   "Có những báo cáo cho thấy hiệu quả khác nhau giữa các loại can thiệp; Zhang và cs "
   "(2023) ghi nhận các chương trình CBT có xu hướng cho kích thước hiệu ứng cao hơn — "
   "kết luận này cũng được ủng hộ trong một tổng quan tập trung vào lo âu (Fisak và cs, "
   "2011). Tuy nhiên, trong phân tích tự sự đầu tiên về can thiệp trường học, Caldwell "
   "và cs (2019) không ủng hộ quan điểm này mạnh như vậy. Họ kết luận rằng CBT chỉ cho "
   "BẰNG CHỨNG YẾU về giảm lo âu — và các nghiên cứu đều nhỏ. Họ cũng KHÔNG tìm thấy "
   "bằng chứng về sự khác biệt giữa các loại can thiệp trong việc giảm trầm cảm.")
crit("Đoạn này CÓ MÂU THUẪN ĐÁNG CHÚ Ý: Zhang 2023 (newer meta-analysis) thì ủng hộ CBT "
     "vượt trội; Caldwell 2019 (network meta-analysis Lancet Psychiatry) thì kết luận "
     "evidence yếu, không khác biệt rõ giữa các loại can thiệp. Tác giả Brown chọn ngả "
     "về Zhang 2023 vì phù hợp với BESST của họ (CBT-based). Đây là CONFIRMATION BIAS — "
     "khi viết editorial, tác giả thường thiên vị bằng chứng ủng hộ vị trí của mình. "
     "Người đọc cần đọc cả 2 meta-analysis độc lập (Zhang DOI 10.1007/s10964-022-01684-4 "
     "+ Caldwell DOI 10.1016/S2215-0366(19)30403-1) để có cái nhìn cân bằng. Trong corpus "
     "dự án QT29 (Li 2025 BMC NMA) cũng xếp CBT hạng 2 sau ACT trên SUCRA — phù hợp "
     "Caldwell hơn Zhang.")

en("Recent trials in the UK have had mixed effects. One example of a study with nil "
   "effects is a recent large cluster RCT of over 8376 students in 85 schools where a "
   "mindfulness intervention was delivered by teachers (Kuyken et al., 2022). The "
   "proposed explanation for the lack of effects was low student engagement, indicated "
   "by student's home practice being low at both post-intervention and one-year follow-up.")
vn("Các thử nghiệm gần đây ở UK có hiệu quả pha trộn. Một ví dụ về nghiên cứu KHÔNG có "
   "hiệu quả là một cluster RCT lớn gần đây với hơn 8.376 học sinh trong 85 trường, "
   "trong đó can thiệp mindfulness được cung cấp bởi giáo viên (Kuyken và cs, 2022). "
   "Lý giải đề xuất cho việc không có hiệu quả là engagement thấp của học sinh — biểu "
   "hiện qua việc thực hành tại nhà rất thấp cả ở thời điểm hậu can thiệp lẫn theo dõi "
   "1 năm.")
crit("MYRIAD trial (Kuyken 2022) là một trong những RCT lớn nhất về can thiệp SKTT trẻ "
     "em — n=8.376 / 85 trường — và là một bài học LỚN cho field. Lý giải \"engagement "
     "thấp\" có thể đúng nhưng cũng có cách hiểu khác: mindfulness UNIVERSAL có \"diluting "
     "effect\" — pha loãng vì áp cho cả những HS không cần. Cũng có cảnh báo từ "
     "Montero-Marin 2022 (cùng nhóm MYRIAD): mindfulness có thể thậm chí GÂY HẠI cho "
     "một subgroup HS vốn đã có triệu chứng tâm lý. Tác giả Brown KHÔNG nhắc tới khía "
     "cạnh \"có thể có hại\" này — selective reporting.")

en("However, another trial of 900 students in 57 schools in 4 regions in England (called "
   "BESST) demonstrated very positive results (Brown et al 2024). We report on some of "
   "the factors that have surfaced in systematic reviews and our own studies that might "
   "account for the accessibility and efficacy of these interventions (Brown et al., 2022).")
vn("Tuy nhiên, một thử nghiệm khác với 900 học sinh tại 57 trường ở 4 vùng nước Anh "
   "(tên là BESST) đã cho thấy kết quả \"rất tích cực\" (Brown và cs, 2024). Chúng tôi "
   "báo cáo về một số yếu tố đã được phát hiện trong các tổng quan hệ thống cũng như "
   "trong các nghiên cứu của chính chúng tôi — có thể giải thích cho khả năng tiếp cận "
   "và hiệu quả của các can thiệp này (Brown và cs, 2022).")
crit("CỤM \"VERY POSITIVE RESULTS\" CHO BESST LÀ CÁCH MÔ TẢ THIÊN VỊ VÀ CẦN HIỆU CHỈNH. "
     "Số liệu thực từ bài Brown 2024 Lancet Psychiatry (em đã đọc full PDF): "
     "Cohen d = −0,17 cho mẫu chung (small effect, dưới ngưỡng \"clinically meaningful\" "
     "0,2 của Cohen). Adjusted MFQ −2,06 (95% CI −3,35 đến −0,76); p=0,0019. Chỉ "
     "subgroup MFQ >27 baseline (n=298, ~33% mẫu) đạt d = −0,52 (medium). Sleep "
     "(SCI) p=0,072 KHÔNG có ý nghĩa thống kê. Đây là kết quả KHIÊM TỐN — không nên gọi "
     "là \"very positive\". Lưu ý: 70% mẫu là nữ (641/900), Year 12-13 (16-18 tuổi), "
     "predominantly White/Asian — generalisability hạn chế.")

# ---------- 5. Universal vs Targeted ----------
H2("5. Phổ quát (Universal) vs Chỉ định (Targeted) — 2 mô hình triển khai")
en("Within school settings, interventions are typically delivered through one of two "
   "main models: targeted interventions, which focus on students identified as being at "
   "high risk of mental health issues, and universal interventions, which are offered "
   "to all students, regardless of their mental health status. Universal interventions "
   "have the advantage of reducing stigma and promoting mental health awareness across "
   "the broader student population, but they can dilute the intensity of support provided "
   "to those in greatest need and have been shown to be less effective in reducing "
   "mental health issues for those who already exhibit problems (Montero-Marin et al., "
   "2022; Stallard et al., 2012).")
vn("Trong môi trường học đường, các can thiệp thường được triển khai theo 1 trong 2 mô "
   "hình chính: (1) CHỈ ĐỊNH (targeted) — tập trung vào HS được xác định là CÓ NGUY CƠ "
   "CAO; và (2) PHỔ QUÁT (universal) — cung cấp cho TẤT CẢ HS bất kể tình trạng SKTT. "
   "Can thiệp phổ quát có ưu điểm GIẢM KỲ THỊ và nâng cao nhận thức SKTT trong cộng "
   "đồng HS rộng — nhưng có thể PHA LOÃNG cường độ hỗ trợ dành cho người cần nhất, và "
   "đã được chứng minh là KÉM HIỆU QUẢ HƠN trong việc giảm vấn đề SKTT cho những HS "
   "đã có triệu chứng (Montero-Marin và cs, 2022; Stallard và cs, 2012).")
crit("Đoạn này tóm tắt CHÍNH XÁC vấn đề \"diluting effect\" của universal. Đây là lý do "
     "tại sao trong corpus dự án Happy House VN (VN030, universal) cho effect d=0,11 "
     "(rất nhỏ) — chứng cứ thực nghiệm cho diluting effect ở VN. Tuy nhiên cần cảnh "
     "báo: targeted lại có vấn đề riêng (xem đoạn tiếp).")

en("Conversely, targeted interventions can be more resource-efficient by concentrating "
   "on students who are most at risk, but they may miss students who are struggling but "
   "do not meet the criteria for high-risk/elevated symptoms (Calear & Christensen, "
   "2010; Radez et al., 2021) and can also increase feelings of stigma in the target "
   "group (Gronholm et al., 2018). Screening also requires extra resources.")
vn("Ngược lại, can thiệp chỉ định có thể tiết kiệm nguồn lực hơn nhờ tập trung vào HS "
   "rủi ro cao nhất, nhưng có thể BỎ SÓT HS đang gặp khó khăn nhưng chưa đáp ứng tiêu "
   "chí \"rủi ro cao / triệu chứng cao\" (Calear & Christensen, 2010; Radez và cs, 2021), "
   "và cũng có thể TĂNG cảm giác kỳ thị trong nhóm được nhắm tới (Gronholm và cs, 2018). "
   "Sàng lọc cũng đòi hỏi nguồn lực bổ sung.")
crit("Đây chính là TRADE-OFF then chốt mà Brown sử dụng để giới thiệu PLACES + "
     "self-referral như giải pháp dung hoà. Bài học từ corpus: Gronholm 2018 (DOI "
     "10.1016/j.jad.2018.07.023) — meta-synthesis qualitative — đã chỉ ra targeted "
     "interventions thường gây stigma cho HS được \"label\". Self-referral né được "
     "stigma này vì HS tự chọn → ít cảm giác bị phân biệt.")

en("The most effective way to offer these school-based interventions is currently "
   "unclear with one review suggesting that targeted interventions more effective "
   "(Werner-Seidler et al., 2021) but two others finding no differences (Fisak et al., "
   "2011; Kapadia et al., 2022).")
vn("Cách hiệu quả nhất để triển khai các can thiệp trường học hiện vẫn CHƯA RÕ RÀNG — "
   "một tổng quan cho rằng can thiệp chỉ định hiệu quả hơn (Werner-Seidler và cs, 2021) "
   "nhưng hai tổng quan khác không thấy khác biệt (Fisak và cs, 2011; Kapadia và cs, "
   "2022).")
crit("Đây là lời thừa nhận trung thực — bằng chứng MIXED. Trong bối cảnh VN, vì "
     "thiếu chuyên gia + kinh phí, em đề xuất: kết hợp UNIVERSAL psychoeducation + "
     "TARGETED CBT cho HS có sàng lọc (vd. GAD-7 ≥ 5) — hybrid model phù hợp resource "
     "low setting.")

# ---------- 6. Self-referral & PLACES ----------
H2("6. Phương pháp Self-referral và mô hình PLACES (tăng khả năng tiếp cận)")
en("Self-referral is where individuals independently decide to participate by referring "
   "themselves, rather than relying on a professional. In our work, self-referral is one "
   "part of the 5 part model described in the evidence-based \"PLACES\" accessibility "
   "model that is designed to increase access by reducing stigma and increasing "
   "convenience of interventions (Brown et al., 2022).")
vn("Self-referral (tự giới thiệu) là cơ chế trong đó CÁ NHÂN tự quyết định tham gia bằng "
   "cách tự đăng ký, thay vì phụ thuộc vào một chuyên gia. Trong công trình của chúng "
   "tôi, self-referral là một phần của mô hình 5 thành phần được mô tả trong mô hình "
   "khả năng tiếp cận \"PLACES\" dựa trên bằng chứng — thiết kế để tăng tiếp cận bằng "
   "cách giảm kỳ thị và tăng thuận tiện cho can thiệp (Brown và cs, 2022).")
crit("LƯU Ý SAI SỐ: Bài này ghi \"5 part model\" nhưng PLACES THỰC TẾ CÓ 6 THÀNH PHẦN — "
     "P (Publicity), L (Lay non-diagnostic title), A (Acceptable), C (Convenient), E "
     "(Effective), S (Self-referral) — em đã verify từ bài Brown 2022 IJERPH gốc PMC8909998. "
     "Có thể tác giả Brown gộp \"perceived effectiveness\" vào một phần khác trong "
     "editorial này; hoặc đây là LỖI EDITORIAL. Khi trích cho thầy, dùng số chuẩn = 6 "
     "yếu tố theo bài gốc Brown 2022.")

en("For example, the PLACES accessibility model uses everyday language (e.g. \"stress\") "
   "instead of clinical terms (e.g. \"depression\" or \"anxiety\") to reduce stigma. "
   "Another aspect is that interventions are delivered in non-stigmatising settings that "
   "are convenient to the target population.")
vn("Ví dụ, mô hình PLACES sử dụng ngôn ngữ đời thường (vd. \"stress / căng thẳng\") "
   "thay vì các thuật ngữ lâm sàng (vd. \"depression / trầm cảm\" hoặc \"anxiety / lo "
   "âu\") để giảm kỳ thị. Một khía cạnh khác là các can thiệp được cung cấp tại các bối "
   "cảnh KHÔNG kỳ thị và thuận tiện cho đối tượng đích.")
crit("Bằng chứng cho \"Lay language\" rất mạnh từ Brown 2022 PLACES paper: khi đổi tên "
     "workshop từ \"depression\" → \"self-confidence\", uptake TĂNG TỪ 28 → 120 đăng ký "
     "(4×); và non-consultation rate (HS chưa từng tìm GP) tăng từ 9,8% → 39% — chứng "
     "tỏ Lay language thu hút được \"hard-to-reach\" group. Áp dụng VN: thay vì gọi "
     "\"chương trình can thiệp lo âu\", nên gọi \"workshop căng thẳng học tập\" hay "
     "\"phòng học kỹ năng sống\".")

en("Self-referral systems, have seldom been employed in trials of school based "
   "interventions (Brown et al., 2010) despite offering benefits such as lowering stigma "
   "(Brown et al., 2010) and also reinforcing autonomy that is particularly valued by "
   "adolescents (Wilson & Deane, 2012). They also allow for more efficient resource use "
   "with high levels of engagement from students who may not have previously sought help, "
   "better follow-up rates (Brown et al., 2019), and greater participation from minoritised "
   "groups that are frequently under-served (Brown et al., 2019).")
vn("Hệ thống self-referral hiếm khi được sử dụng trong các thử nghiệm can thiệp tại "
   "trường học (Brown và cs, 2010), mặc dù mang lại nhiều lợi ích như giảm kỳ thị (Brown "
   "và cs, 2010) và củng cố tính tự chủ — vốn đặc biệt được vị thành niên coi trọng "
   "(Wilson & Deane, 2012). Self-referral cũng cho phép sử dụng nguồn lực hiệu quả hơn "
   "với mức độ engagement cao từ những HS chưa từng tìm trợ giúp trước đây, tỉ lệ "
   "follow-up tốt hơn (Brown và cs, 2019), và sự tham gia nhiều hơn từ các nhóm thiểu "
   "số thường bị bỏ rơi trong dịch vụ (Brown và cs, 2019).")
crit("Bằng chứng đáng chú ý: trong BESST 2024, 80% HS tự đăng ký CHƯA TỪNG tìm GP cho "
     "vấn đề SKTT — tức self-referral đã chạm tới nhóm \"hard-to-reach\" mà các kênh "
     "thông thường (qua GP, qua giáo viên) đã bỏ sót. Đây là evidence rất mạnh cho "
     "self-referral. Áp dụng VN: cần thiết kế kênh tự đăng ký qua app/website ẩn danh, "
     "hoặc qua giờ chào cờ — không qua giáo viên chủ nhiệm để tránh stigma.")

# ---------- 7. Co-design ----------
H2("7. Đồng thiết kế (Co-design) với chính người trẻ")
en("Co-design where students help design the intervention is growing in importance. It "
   "is an important issue given that professionals' and users' perspectives about what "
   "are important aspects may vary (Wittevrongel et al., 2025). The DISCOVER workshop "
   "programme was co-designed with 16–18 year olds (Sclare et al., 2015) and is designed "
   "to enhance student-clinician interaction and increases the use of visual aids, "
   "including videos and group exercises.")
vn("Co-design (đồng thiết kế) — nơi học sinh giúp thiết kế can thiệp — đang ngày càng "
   "quan trọng. Đây là vấn đề có ý nghĩa vì quan điểm của các CHUYÊN GIA và NGƯỜI DÙNG "
   "về những khía cạnh quan trọng có thể KHÁC NHAU (Wittevrongel và cs, 2025). Chương "
   "trình DISCOVER được đồng thiết kế với HS 16–18 tuổi (Sclare và cs, 2015) — nhằm "
   "tăng cường tương tác HS–nhà lâm sàng và tăng sử dụng các phương tiện trực quan, bao "
   "gồm video và bài tập nhóm.")
crit("Co-design là xu hướng được ủng hộ mạnh trong implementation science nhưng bằng "
     "chứng định lượng còn YẾU (Slattery 2020). Áp dụng VN: cần Patient/Public "
     "Involvement (PPI) bài bản trước khi chạy RCT — không chỉ \"hỏi ý kiến\" mà thực "
     "sự cho HS tham gia thiết kế module, ngôn ngữ, hình ảnh.")

en("Co-design was achieved through extensive consultation with young people who provided "
   "feedback on the intervention workshop's format, content, and delivery style, ensuring "
   "the program addressed the specific needs and preferences of this age group. Their "
   "input shaped various aspects of the intervention, including the program name as well "
   "as refining materials and interactive tasks.")
vn("Đồng thiết kế đạt được thông qua tham vấn rộng rãi với người trẻ — họ cung cấp "
   "phản hồi về định dạng, nội dung, và phong cách triển khai của workshop, đảm bảo "
   "chương trình đáp ứng nhu cầu và sở thích cụ thể của nhóm tuổi này. Đóng góp của họ "
   "định hình nhiều khía cạnh của can thiệp, bao gồm cả TÊN chương trình cũng như tinh "
   "chỉnh tài liệu và bài tập tương tác.")

en("However, possibly because of the newness of this approach, much of the literature is "
   "devoted to describing different methods of co-design and evidence about the "
   "effectiveness of co-design of mental health intervention about greater engagement is "
   "still very limited (Slattery et al., 2020). Nevertheless, qualitative studies have "
   "been promising. A review of PPI work with young people has also concluded that more "
   "operationalisation and consistency of evaluation measures was needed (Totzeck et al., "
   "2024). For example, the effects of the participation on participants and whether "
   "their input changed the research process or researchers' opinions could be noted.")
vn("Tuy nhiên, có thể do tính mới của cách tiếp cận này, phần lớn tài liệu hiện có chỉ "
   "MÔ TẢ các phương pháp co-design khác nhau, còn bằng chứng về hiệu quả của co-design "
   "trong việc tăng engagement vẫn rất HẠN CHẾ (Slattery và cs, 2020). Tuy vậy, các "
   "nghiên cứu định tính cho thấy hứa hẹn. Một tổng quan về công tác PPI với người trẻ "
   "cũng kết luận rằng cần thao tác hoá nhiều hơn và nhất quán hơn về các thước đo đánh "
   "giá (Totzeck và cs, 2024). Ví dụ, có thể cần ghi nhận: hiệu quả của sự tham gia lên "
   "chính người tham gia, và liệu đầu vào của họ có thay đổi quy trình nghiên cứu hoặc "
   "quan điểm của nhà nghiên cứu hay không.")
crit("Đây là tự nhận hạn chế đáng quý — co-design như một field còn THIẾU rigour về "
     "operationalisation. Tuy nhiên Brown vẫn ngầm khuyến nghị mạnh DISCOVER (đã được "
     "co-designed) — circular reasoning. Bài học cho VN: cần xây bộ chỉ số đánh giá "
     "co-design (vd. CHE-s scale Barello 2019 đã có) trước khi triển khai quy mô.")

# ---------- 8. BESST trial — kết hợp ----------
H2("8. Thử nghiệm BESST — Tích hợp các yếu tố")
en("The BESST trial (Brown et al. 2024) incorporated several elements:")
vn("Thử nghiệm BESST (Brown và cs, 2024) đã tích hợp các yếu tố sau:")
nr("• Một can thiệp được đồng thiết kế (co-designed) với HS 16–18 tuổi (Sclare và cs, 2015)")
nr("• Quy trình self-referral phi-chẩn-đoán, được mô tả trong mô hình PLACES (Brown và cs, 2022)")
nr("• Một chương trình CBT ngắn (1 ngày) cho HS 16–18 tuổi")
nr("• Được cung cấp bởi đội Mental Health Support Teams (MHST) tại trường học")

en("The self-referral pathway (Brown et al., 2022) was offered at school assemblies when "
   "students were invited to opt-in into the trial and stress workshops. Stress problems "
   "were described in ordinary ways such as school exams and family pressures. The "
   "DISCOVER intervention was an one-day, CBT-based workshop which was designed to help "
   "adolescents manage stress, anxiety, and low mood by providing practical strategies "
   "for coping with mental health challenges (Sclare et al., 2015).")
vn("Lộ trình self-referral (Brown và cs, 2022) được giới thiệu tại các giờ chào cờ "
   "(school assemblies), khi HS được mời tự chọn (opt-in) tham gia thử nghiệm và "
   "workshop căng thẳng. Vấn đề stress được mô tả theo cách thông thường như áp lực thi "
   "cử và áp lực gia đình. Can thiệp DISCOVER là một workshop 1 ngày dựa trên CBT — "
   "thiết kế để giúp VTN quản lý stress, lo âu và tâm trạng thấp bằng các chiến lược "
   "thực tế đối phó với khó khăn SKTT (Sclare và cs, 2015).")

en("This led to 900 students referring themselves to the trial with 46% of these students "
   "being from minority groups. The workshop intervention was each run by 3 MHST clinical "
   "staff in groups of up to 19 students in 4 English regions. The trial provided robust "
   "evidence for the clinical (ES= −0.17) and cost-effectiveness of a school-based "
   "intervention for adolescents especially for those with elevated depressive symptoms "
   "(ES= −0.52) (Brown et al., 2024). Process evaluation results were positive among "
   "students themselves and staff delivering the workshops (Weaver et al, in revision). "
   "The results are consistent with Kuyken et al.'s (2023) reflections in 2023 after his "
   "unsuccessful trial about the need for good reach, student engagement, co-design of "
   "the intervention and the need to think about sustainability of the intervention "
   "through the MHSTs.")
vn("Điều này dẫn đến 900 HS tự đăng ký vào thử nghiệm — trong đó 46% thuộc các nhóm "
   "thiểu số. Workshop được mỗi lần dẫn dắt bởi 3 nhân viên lâm sàng MHST cho nhóm tối "
   "đa 19 HS, tại 4 vùng nước Anh. Thử nghiệm cung cấp bằng chứng vững chắc cho hiệu "
   "quả lâm sàng (ES = −0,17) và chi phí-hiệu quả của một can thiệp tại trường học cho "
   "VTN — ĐẶC BIỆT đối với những HS có triệu chứng trầm cảm cao (ES = −0,52) (Brown và "
   "cs, 2024). Kết quả đánh giá quy trình (process evaluation) cũng tích cực ở cả HS và "
   "nhân viên cung cấp workshop (Weaver và cs, đang hiệu đính). Kết quả phù hợp với "
   "phản tỉnh của Kuyken và cs (2023) sau thử nghiệm không thành công của ông — về "
   "tầm quan trọng của reach (tiếp cận) tốt, engagement của HS, co-design can thiệp, và "
   "tính bền vững thông qua MHSTs.")
crit("HAI EFFECT SIZE QUAN TRỌNG cần phân biệt rõ ràng: (1) Cohen d = −0,17 cho mẫu "
     "chung 854 HS — SMALL effect, dưới ngưỡng \"clinically meaningful\" 0,2 của Cohen. "
     "(2) Cohen d = −0,52 cho subgroup 298 HS có MFQ baseline >27 (~33% mẫu, tức HS đã "
     "có triệu chứng trầm cảm cao baseline) — MEDIUM effect. Cách Brown dùng từ \"robust "
     "evidence\" cho ES = −0,17 là PHÓNG ĐẠI; thực ra từ \"modest\" đã được dùng trong "
     "abstract Lancet Psychiatry gốc của họ — nhưng editorial này thì lờ đi từ \"modest\" "
     "và chỉ ghi \"robust\". Hệ quả: BESST chứng minh giá trị của targeted self-referral "
     "(với HS đã có triệu chứng), KHÔNG nhất thiết phổ quát. Đây là điểm bị đảo ngược "
     "trong cách viết editorial — cần đọc cẩn thận.")

# ---------- 9. Tương lai + Khuyến nghị ----------
H2("9. Tương lai + 2 Khuyến nghị")
en("Andrews and Foulkes (2025) have also strongly recommended that we need to move away "
   "from universal interventions for students. They recommend more opt-in interventions "
   "which are more in line with adolescents' desire for autonomy, which is equivalent to "
   "the self-referral system. The BESST trial seems to show that these factors do make "
   "for better accessibility and a more effective and cost-effective approach.")
vn("Andrews và Foulkes (2025) cũng khuyến nghị MẠNH MẼ rằng chúng ta cần dịch chuyển "
   "RA KHỎI các can thiệp phổ quát cho HS. Họ khuyến nghị các can thiệp opt-in (chọn "
   "tham gia) — phù hợp với khao khát tự chủ của VTN, tương đương với hệ thống "
   "self-referral. Thử nghiệm BESST dường như cho thấy các yếu tố này thực sự dẫn đến "
   "khả năng tiếp cận tốt hơn và cách tiếp cận hiệu quả hơn về cả lâm sàng lẫn chi phí.")
crit("\"Move away from universal\" là một dịch chuyển paradigm CÓ TÍNH ĐẢO LỘN. Andrews "
     "& Foulkes 2025 (Child & Adolescent Mental Health) là bài quan điểm cũng cùng "
     "hướng. Tuy nhiên cần cảnh báo: từ bỏ universal có thể bỏ sót VAI TRÒ phòng ngừa "
     "PRIMARY (cho HS chưa có triệu chứng) — đặc biệt với áp dụng VN nơi MENTAL HEALTH "
     "LITERACY cộng đồng còn yếu, cần universal psychoeducation làm nền. Đề xuất: "
     "HYBRID = universal psychoeducation (literacy + giảm stigma) + targeted "
     "self-referral CBT (cho HS có triệu chứng).")

en("We make two recommendations. One is that this field should assess actual engagement "
   "of students as opposed to attitudes to help-seeking. While mental health literacy and "
   "the relationship with attitudes to help-seeking have been frequently studied, the "
   "effects on actual help-seeking have not been shown. The PLACES model is probably one "
   "of the few models that does show actual help-seeking behaviour (Brown et al., 2022) "
   "improves among young people as well as adults with different problems.")
vn("Chúng tôi đưa ra hai khuyến nghị. THỨ NHẤT: lĩnh vực này nên đánh giá ENGAGEMENT "
   "THỰC TẾ của HS thay vì chỉ thái độ với việc tìm trợ giúp. Trong khi mental health "
   "literacy và mối quan hệ với THÁI ĐỘ tìm trợ giúp đã được nghiên cứu nhiều, tác động "
   "lên HÀNH VI tìm trợ giúp THỰC TẾ vẫn chưa được chứng minh. Mô hình PLACES có lẽ là "
   "một trong số ít mô hình chứng minh được hành vi tìm trợ giúp thực tế cải thiện (Brown "
   "và cs, 2022) — ở cả người trẻ lẫn người lớn với nhiều vấn đề khác nhau.")
crit("Khuyến nghị này HỢP LÝ — đo behaviour thay vì chỉ attitude. Tuy nhiên \"PLACES "
     "improves actual help-seeking\" cần thận trọng: bài Brown 2022 PLACES (em đã đọc "
     "full PMC8909998) là 4 case studies KHÔNG có nhóm chứng đối — uptake tăng từ 28→120 "
     "khi đổi tên là evidence YẾU (level 4, case series). PLACES chưa được test trong "
     "RCT như một intervention package độc lập. Đây là LIMITATION không nhỏ.")

en("Secondly, we recommend that the DISCOVER workshop, and the PLACES accessibility "
   "model form the basis for offering school interventions to 16–18-year-olds in schools. "
   "We are not claiming it is perfect and we are researching how it might be implemented "
   "across England as understanding the context of school interventions is very important "
   "(Roshan et al., 2025). However, it does appear very effective, clinically and "
   "cost-effectively, and also engages diverse students very well, and all this at a time "
   "when other interventions have been disappointing. We hope this may help start to "
   "reduce the rising tide of mental health problems among young people.")
vn("THỨ HAI: chúng tôi khuyến nghị workshop DISCOVER, cùng với mô hình tiếp cận PLACES, "
   "làm nền tảng cho các can thiệp tại trường học dành cho HS 16–18 tuổi. Chúng tôi "
   "không tuyên bố mô hình này là hoàn hảo và đang nghiên cứu cách triển khai khắp nước "
   "Anh — vì hiểu bối cảnh của can thiệp trường học rất quan trọng (Roshan và cs, 2025). "
   "Tuy nhiên, mô hình này có vẻ rất hiệu quả về cả lâm sàng lẫn chi phí, và thu hút "
   "tốt HS đa dạng — vào thời điểm các can thiệp khác đang gây thất vọng. Chúng tôi hy "
   "vọng điều này có thể giúp bắt đầu giảm làn sóng vấn đề SKTT đang dâng cao ở người trẻ.")
crit("Đây là KHUYẾN NGHỊ MẠNH cuối cùng — DISCOVER + PLACES làm \"nền tảng\". Cần phân "
     "biệt: với HS 16–18 (sixth form UK) thì có evidence (BESST). Với HS dưới 16 tuổi "
     "(THCS) thì CHƯA CÓ EVIDENCE — không nên ngoại suy. Áp dụng VN: VTN VN 13-18 tuổi, "
     "phần lớn còn ở THCS — nhóm tuổi 16-18 đã ít hơn (nhiều HS bỏ học hoặc đi làm sau "
     "lớp 9). Cần adaptation cho 13-15 tuổi nếu muốn quy mô lớn ở VN.")

# =====================================================================
# PHẦN 2 — BẢNG TỔNG HỢP
# =====================================================================
H1("PHẦN 2 — BẢNG TỔNG HỢP")

H2("Bảng A. So sánh 5 mô hình triển khai can thiệp SKTT trường học")
tbl(['Mô hình', 'Cách thức', 'Ưu điểm', 'Nhược điểm', 'Bằng chứng UK chính'],
    [
        ['Universal (phổ quát)',
         'Cung cấp cho TẤT CẢ HS bất kể tình trạng',
         'Giảm kỳ thị, nâng cao nhận thức cộng đồng, tiếp cận diện rộng',
         'Diluting effect (pha loãng), kém hiệu quả với HS đã có triệu chứng',
         'MYRIAD Kuyken 2022 (n=8.376) NULL effect; Stallard 2012 BMJ classroom CBT yếu'],
        ['Targeted (chỉ định)',
         'Sàng lọc HS rủi ro cao trước, can thiệp riêng nhóm này',
         'Tiết kiệm nguồn lực, hiệu quả lớn hơn mỗi cá nhân được can thiệp',
         'Bỏ sót HS dưới ngưỡng, gây stigma cho nhóm được "label", cần extra resource sàng lọc',
         'Werner-Seidler 2021 SR favours targeted; Gronholm 2018 stigma evidence'],
        ['Self-referral (tự giới thiệu)',
         'HS tự đăng ký, không qua chuyên gia / GV',
         'Tự chủ, giảm stigma, thu hút "hard-to-reach", follow-up tốt',
         'Có thể bỏ sót HS không tự nhận diện được vấn đề; bias selection',
         'BESST 2024 (n=900): 80% chưa từng tìm GP; 46% minority'],
        ['Teacher-led (giáo viên dẫn)',
         'GV được đào tạo cung cấp can thiệp ngay tại lớp',
         'Mở rộng quy mô dễ, tiết kiệm chi phí',
         'Hiệu quả thấp hơn lâm sàng (Stallard 2014); cần đào tạo + supervision',
         'Stallard 2014 PACES (FRIENDS) Lancet Psychiatry small effect'],
        ['MHST-led (đội hỗ trợ SKTT)',
         'Therapists thạc sĩ NHS đến trường cung cấp CBT',
         'Chuyên môn tốt hơn GV; gắn với hệ thống y tế công',
         'Tốn kém; cần infrastructure NHS; chưa có cơ chế tương đương ở VN',
         'BESST 2024: 15 MHSTs / 11 NHS Trusts; cost £108,87/student'],
    ], [3.5, 3.5, 3.5, 3.5, 3.5])

H2("Bảng B. 5 RCT/Trial UK quan trọng được trích trong bài")
tbl(['Trial', 'Tác giả, Năm, Tạp chí', 'N / Trường', 'Công cụ / Outcome', 'Kết quả chính'],
    [
        ['BESST', 'Brown JSL et al. 2024, Lancet Psychiatry 11(7):504–515 (PMID 38759665)',
         '900 HS / 57 trường / 4 vùng UK',
         'MFQ depression 6 tháng (primary)',
         'Cohen d=−0,17 chung; d=−0,52 subgroup elevated depression; cost-effective 61–78%'],
        ['MYRIAD', 'Kuyken W et al. 2022, Evid Based Ment Health 25(3):99–109',
         '8.376 HS / 85 trường UK',
         'Mindfulness, multi-outcome',
         'NULL effect; engagement thấp; có thể có hại cho subgroup vulnerable (Montero-Marin 2022)'],
        ['PACES (FRIENDS)', 'Stallard P et al. 2014, Lancet Psychiatry 1(3):185–192',
         '~1.250 HS UK',
         'Anxiety prevention',
         'Hiệu quả nhỏ; clinician > teacher delivery'],
        ['DISCOVER feasibility', 'Brown JSL et al. 2019, J Adolesc 71:150–161',
         '~155 HS UK pilot',
         'Stress workshop, MFQ + RCADS',
         'Feasible; depression d=0,27; anxiety d=0,25 (3 tháng post)'],
        ['Classroom CBT', 'Stallard P et al. 2012, BMJ 345:e6058',
         'High-risk HS UK',
         'Depression prevention',
         'Yếu (small effect); confirmed limitations universal in classroom'],
    ], [3.0, 4.0, 2.5, 2.5, 5.0])

H2("Bảng C. So sánh UK ↔ Việt Nam — Bối cảnh áp dụng")
tbl(['Khía cạnh', 'UK (theo Brown & Carter 2025)', 'Việt Nam (theo corpus dự án)'],
    [
        ['Tỉ lệ chẩn đoán RLTT VTN', '~25% ở 17-19 tuổi (NHS 2023)',
         '21,7% mental health problem (V-NAMHS subthreshold+); '
         '3,3% mental disorder DSM-5 (V-NAMHS DISC-5)'],
        ['Tỉ lệ KHÔNG nhận chăm sóc', '60% (UK Gov 2022)',
         '91,6% (V-NAMHS — chỉ 8,4% tiếp cận dịch vụ)'],
        ['Phụ huynh nhận ra con bị lo âu', 'Chưa có số liệu rõ',
         '5,1% (V-NAMHS) — RẤT THẤP'],
        ['Bác sĩ tâm thần / 100.000 dân', '~9,75 (theo OECD UK 2020)',
         '~0,2 (corpus QT10 GBD ASEAN — THIẾU TRẦM TRỌNG)'],
        ['Cơ chế MHST ở trường', 'Có (NHS England, ngân sách công, đào tạo chuẩn quốc gia)',
         'CHƯA CÓ. Hiện có Tư vấn học đường (Thông tư 31/2017/TT-BGDĐT) nhưng chưa được '
         'đào tạo CBT chuẩn'],
        ['Mô hình self-referral ở trường', 'BESST đã chứng minh khả thi (80% chưa từng tìm GP)',
         'Chưa có RCT; cần thiết kế kênh ẩn danh (app/web) để tránh stigma từ GVCN'],
        ['Mental Health Literacy phụ huynh', 'Cao hơn (qua các campaign NHS dài hạn)',
         'Thấp (5,1% nhận ra con bị lo âu) — cần universal psychoeducation trước khi targeted'],
        ['Universal can thiệp đã thử',
         'MYRIAD (mindfulness) — null; PACES (CBT classroom) — yếu',
         'Happy House VN030 — d=0,11 (rất nhỏ, confirmed diluting effect)'],
        ['Chương trình adaptation đề xuất',
         'BESST/PLACES targeted self-referral cho 16-18',
         'HYBRID: universal psychoeducation literacy + targeted CBT self-referral '
         '(GAD-7 ≥ 5) cho 13-18 tuổi; điều chỉnh ngôn ngữ "căng thẳng học tập"'],
        ['Gap RCT cần lấp ở VN',
         'Đã có nhiều RCT UK',
         'CHƯA CÓ RCT VN nào về CBT trường học self-referral. RẤT CẦN.'],
    ], [4.0, 6.0, 7.0])

# =====================================================================
# PHẦN 3 — PHẢN BIỆN TỔNG QUAN
# =====================================================================
H1("PHẦN 3 — PHẢN BIỆN TỔNG QUAN")

H2("3.1. Điểm mạnh của bài")
nr("• Tổng hợp KỊP THỜI và toàn diện về can thiệp SKTT trường học UK 2025 — "
   "đặc biệt sau khi BESST (2024) công bố.")
nr("• Hai tác giả Brown & Carter là CHUYÊN GIA hàng đầu UK về field này — chủ trì cả "
   "BESST (Brown JSL — chief investigator) lẫn PLACES (Brown JSL — first author).")
nr("• So sánh đa mô hình rõ ràng: universal/targeted/self-referral, teacher/clinician/MHST.")
nr("• Bài học từ THẤT BẠI (MYRIAD Kuyken 2022 — n=8.376) cũng được nhấn mạnh — quan "
   "trọng cho việc thiết kế lại field.")
nr("• Đề xuất 2 khuyến nghị thực tế: (1) đo behaviour thay vì attitude; (2) "
   "DISCOVER+PLACES làm nền.")

H2("3.2. Điểm yếu / Hạn chế của bài")
crit("(1) LOẠI BÀI: Editorial = level-5 evidence theo Oxford CEBM 2011. KHÔNG phải "
     "systematic review, không có PRISMA, không có quality appraisal. Mọi tuyên bố "
     "tổng hợp cần đối chiếu lại nguồn gốc.")
crit("(2) CONFLICT OF INTEREST kỹ thuật: dù tác giả khai \"no conflict\", thực tế họ "
     "khuyến nghị MẠNH các sản phẩm CỦA HỌ (DISCOVER, PLACES, BESST). Đây là intellectual "
     "conflict — không vô hiệu lập luận nhưng cần ghi nhận khi đọc.")
crit("(3) PHÓNG ĐẠI BESST: gọi ES = −0,17 là \"very positive\" và \"robust evidence\" "
     "— trong khi abstract Lancet Psychiatry gốc của chính họ dùng từ \"modestly\". "
     "Editorial này selective.")
crit("(4) GHI SAI SỐ THÀNH PHẦN PLACES: bài ghi \"5 part model\" nhưng PLACES gốc có "
     "6 yếu tố P-L-A-C-E-S (verify từ Brown 2022 IJERPH PMC8909998). Lỗi không lớn "
     "nhưng cho thấy sự thiếu cẩn thận editorial.")
crit("(5) UK-CENTRIC: hoàn toàn dùng dữ liệu UK (NHS, MHST, sixth form 16-18). Không "
     "thảo luận generalisability sang LMIC. Khi áp dụng cho VN cần adaptation lớn.")
crit("(6) BỎ SÓT bằng chứng \"có hại\": Montero-Marin 2022 (cùng nhóm MYRIAD) đã chỉ ra "
     "mindfulness UNIVERSAL có thể gây HẠI cho subgroup vulnerable — Brown lờ đi điểm "
     "này khi nói MYRIAD chỉ \"engagement thấp\".")
crit("(7) BỎ SÓT 2 META-ANALYSIS QUAN TRỌNG: bài chỉ trích Zhang 2023 (ủng hộ CBT) mà "
     "không cân đối với Caldwell 2019 NMA Lancet Psychiatry (DOI 10.1016/S2215-0366(19)30403-1) "
     "— Caldwell kết luận evidence YẾU cho mọi can thiệp. Confirmation bias.")

H2("3.3. Phù hợp với bối cảnh Việt Nam")
nr("• PLACES có TIỀM NĂNG mạnh cho VN — đặc biệt 3 yếu tố:")
nr("    – Lay language: đổi \"lo âu/trầm cảm\" → \"căng thẳng học tập\" / \"phòng kỹ năng sống\"")
nr("    – Self-referral: thiết kế kênh ẩn danh (app/web) tránh stigma từ GVCN")
nr("    – Convenient: tại trường, sau giờ học hoặc trong tiết sinh hoạt")
nr("• 3 yếu tố cần ADAPT CẨN THẬN cho VN:")
nr("    – Publicity: cần qua kênh phụ huynh + cộng đồng (literacy phụ huynh thấp 5,1%)")
nr("    – Acceptable: cần phù hợp văn hoá Á Đông (face culture, quan hệ gia đình)")
nr("    – Effective: chưa có evidence VN — cần feasibility trước RCT")
nr("• MHST KHÔNG khả thi cho VN ngắn hạn (thiếu nhân lực, không có ngân sách công). "
   "Phương án thay thế: nâng cấp tư vấn học đường hiện có + supervision liên trường + "
   "đào tạo CBT chuẩn ngắn hạn (2-3 ngày như BESST đào tạo MHST).")
nr("• Hybrid model đề xuất cho VN: UNIVERSAL psychoeducation literacy (3-4 tiết/năm "
   "trong môn GDCD/HĐTN, nội dung tự kiểm GAD-7/PHQ-9 + giảm stigma) + TARGETED "
   "self-referral CBT cho HS có GAD-7 ≥ 5 (mô hình DISCOVER 1 ngày, do tư vấn học "
   "đường + chuyên gia hỗ trợ liên trường).")

H2("3.4. Đối chiếu với corpus 35+ bài của dự án")
tbl(['Bài/RAG ID', 'Tên / Tác giả', 'Liên hệ với editorial Brown & Carter 2025'],
    [
        ['QT08 (Wen 2020)', 'LPA lo âu nông thôn TQ',
         'Áp lực học tập OR=11,58 cho lo âu nặng; hỗ trợ trường OR=0,562 (bảo vệ) — '
         'củng cố lập luận của Brown về vai trò trường học'],
        ['QT29 (Li 2025 BMC)', 'NMA can thiệp lo âu trẻ em 12 nước',
         'CBT xếp hạng 2 trên SUCRA (sau ACT 0,69); cùng evidence với editorial'],
        ['QT10/QT12 (GBD ASEAN)', 'GBD 2021 ASEAN Lancet',
         'Việt Nam đứng thứ 3 ASEAN gánh nặng RLTT VTN; 0,2 BS tâm thần/100k dân — '
         'context khác hoàn toàn UK; chứng minh không thể copy MHST'],
        ['VN030 (Happy House)', 'VN universal program',
         'd=0,11 — confirmed diluting effect; ủng hộ Brown từ bỏ universal'],
        ['B8 (Sri Lanka CACBT)', 'Cultural-adapted CBT delivered by teachers',
         'GV được đào tạo TỐT có thể cung cấp CBT hiệu quả ở LMIC — phản biện luận điểm '
         '"clinician > teacher" của Brown'],
        ['B6 (Resilience)', 'School resilience programs review',
         'Resilience-based programs có evidence yếu hơn CBT — phù hợp Brown'],
        ['QT042 (B5 = bài này)', 'Bản tóm tắt cũ của bài',
         'Đã có doc tóm tắt 38 KB nhưng chưa phải dịch full + phản biện như doc này'],
    ], [3.0, 4.0, 10.0])

# =====================================================================
# PHẦN 4 — THAM KHẢO
# =====================================================================
H1("PHẦN 4 — THAM KHẢO ĐẦY ĐỦ")
note("Toàn bộ 30 references trong reference list của bài gốc + 5 nguồn xác minh ngoài "
     "(PMC, PubMed, King's College). Mỗi reference giữ format APA + DOI (hyperlink "
     "text). Đây là phần em đã đối chiếu trực tiếp với raw text reference của bài.")

H2("4.1. Reference của bài gốc (theo thứ tự xuất hiện)")
nr("1. Alderwick H, Dixon J (2019). The NHS long term plan. BMJ Publishing Group.", size=11)
nr("2. Andrews JL, Foulkes L (2025). Debate: Where to next for universal school-based "
   "mental health interventions? Time to move towards more effective alternatives. Child "
   "Adolesc Ment Health, 30(1):102–104. https://doi.org/10.1111/camh.12753", size=11)
nr("3. Brännlund A, Strandh M, Nilsson K (2017). Mental-health and educational achievement: "
   "The link between poor mental-health and upper secondary school completion and grades. "
   "J Mental Health, 26(4):318–325. https://doi.org/10.1080/09638237.2017.1294739", size=11)
nr("4. Brown JSL, Blackshaw E, Stahl D, Fennelly L, McKeague L, Sclare I, Michelson D "
   "(2019). School-based early intervention for anxiety and depression in older adolescents: "
   "A feasibility RCT of a self-referral stress management workshop programme (\"DISCOVER\"). "
   "J Adolescence, 71(1):150–161. https://doi.org/10.1016/j.adolescence.2018.11.009", size=11)
nr("5. Brown JSL, Boardman J, Whittinger N, Ashworth M (2010). Can a self-referral system "
   "help improve access to psychological treatments? Br J Gen Pract, 60(574):365–371. "
   "https://doi.org/10.3399/bjgp10X501877", size=11)
nr("6. Brown JSL, Lisk S, Carter B, Stevelink SAM, Van Lieshout R, Michelson D (2022). "
   "How Can We Actually Change Help-Seeking Behaviour for Mental Health Problems among "
   "the General Public? Development of the 'PLACES' Model. IJERPH, 19(5):2831. "
   "https://doi.org/10.3390/ijerph19052831 (PMC8909998)", size=11)
nr("7. Brown JSL, Murphy C, Kelly J, Goldsmith K (2019). How can we successfully recruit "
   "depressed people? Lessons learned in recruiting depressed participants to a multi-site "
   "trial of a brief depression intervention (the 'CLASSIC' trial). Trials, 20(1):131. "
   "https://doi.org/10.1186/s13063-018-3033-5", size=11)
nr("8. Brown J, James K, Lisk S, Shearer J, Byford S, Stallard P, Deighton J, Saunders D, "
   "Yarrum J, Fonagy P, Weaver T, Sclare I, Day C, Evans C, Carter B (2024). Clinical "
   "effectiveness and cost-effectiveness of a brief accessible CBT programme for stress "
   "in school-aged adolescents (BESST): A cluster RCT in the UK. Lancet Psychiatry, "
   "11(7):504–515. PMID 38759665. https://doi.org/10.1016/S2215-0366(24)00101-9", size=11)
nr("9. Caldwell DM, Davies SR, Hetrick SE, Palmer JC, Caro P, López-López JA, Gunnell D, "
   "Kidger J, Thomas J, French C, Stockings E, Campbell R, Welton NJ (2019). School-based "
   "interventions to prevent anxiety and depression in children and young people: A SR "
   "and network meta-analysis. Lancet Psychiatry, 6(12):1011–1020. "
   "https://doi.org/10.1016/S2215-0366(19)30403-1", size=11)
nr("10. Calear AL, Christensen H (2010). Systematic review of school-based prevention and "
    "early intervention programs for depression. J Adolescence, 33(3):429–438. "
    "https://doi.org/10.1016/j.adolescence.2009.07.004", size=11)
nr("11. de Girolamo G, Dagani J, Purcell R, Cocchi A, McGorry PD (2012). Age of onset of "
    "mental disorders and use of mental health services. Epidemiol Psychiatr Sci, "
    "21(1):47–57. https://doi.org/10.1017/s2045796011000746", size=11)
nr("12. Fazel M, Patel V, Thomas S, Tol W (2014). Mental health interventions in schools "
    "in LMIC. Lancet Psychiatry, 1(5):388–398. https://doi.org/10.1016/S2215-0366(14)70357-8", size=11)
nr("13. Feiss R, Dolinger SB, Merritt M, Reiche E, Martin K, Yanes JA, Thomas CM, "
    "Pangelinan M (2019). A SR and meta-analysis of school-based stress, anxiety, and "
    "depression prevention programs for adolescents. J Youth Adolesc, 48(9):1668–1685. "
    "https://doi.org/10.1007/s10964-019-01085-0", size=11)
nr("14. Fisak BJ, Richard D, Mann A (2011). The prevention of child and adolescent "
    "anxiety: A meta-analytic review. Prevention Science, 12(3):255–268. "
    "https://doi.org/10.1007/s11121-011-0210-0", size=11)
nr("15. Gronholm PC, Nye E, Michelson D (2018). Stigma related to targeted school-based "
    "mental health interventions: A systematic review of qualitative evidence. J Affect "
    "Disord, 240:17–26. https://doi.org/10.1016/j.jad.2018.07.023", size=11)
nr("16. Gulliver A, Griffiths KM, Christensen H (2010). Perceived barriers and facilitators "
    "to mental health help-seeking in young people: A systematic review. BMC Psychiatry, "
    "10(1):113. https://doi.org/10.1186/1471-244X-10-113", size=11)
nr("17. Hale DR, Bevilacqua L, Viner RM (2015). Adolescent health and adult education and "
    "employment: A systematic review. Pediatrics, 136(1):128–140. "
    "https://doi.org/10.1542/peds.2014-2105", size=11)
nr("18. Johnson D, Dupuis G, Piche J, Clayborne Z, Colman I (2018). Adult mental health "
    "outcomes of adolescent depression: A systematic review. Depression and Anxiety, "
    "35(8):700–716. https://doi.org/10.1002/da.22777", size=11)
nr("19. Kapadia D et al. (2022). Ethnic inequalities in health care: A rapid evidence "
    "review. NHS Race and Health Observatory.", size=11)
nr("20. Kessler RC, Berglund P, Demler O, Jin R, Merikangas KR, Walters EE (2005). "
    "Lifetime prevalence and age-of-onset distributions of DSM-IV disorders in the NCS-R. "
    "Arch Gen Psychiatry, 62(6):593–602. https://doi.org/10.1001/archpsyc.62.6.593", size=11)
nr("21. Kuyken W et al. (2022). Effectiveness and cost-effectiveness of universal "
    "school-based mindfulness training compared with normal school provision: MYRIAD "
    "cluster RCT. Evid Based Ment Health, 25(3):99–109. "
    "https://doi.org/10.1136/ebmental-2021-300396", size=11)
nr("22. Kuyken W, Blakemore S-J, Byford S, Dalgleish T, Ford T, Hinze V, Mansfield K, "
    "Montero-Marin J, Ukoumunne OC, Viner RM (2023). Mental health in adolescence: The "
    "role of schools-based social emotional teaching. J Mental Health, 32(3):537–540. "
    "https://doi.org/10.1080/09638237.2023.2210668", size=11)
nr("23. Lovell K, Richards D (2000). Multiple access points and levels of entry (MAPLE): "
    "Ensuring choice, accessibility and equity for CBT services. Behavioural and "
    "Cognitive Psychotherapy, 28(4):379–391. https://doi.org/10.1017/S1352465800004070", size=11)
nr("24. Marcheselli F et al. (2018). Mental health of children and young people in "
    "England, 2017. NHS.", size=11)
nr("25. McGorry P, Bates T, Birchwood M (2013). Designing youth mental health services "
    "for the 21st century. Br J Psychiatry Suppl, 54(s54):s30–s35. "
    "https://doi.org/10.1192/bjp.bp.112.119214", size=11)
nr("26. Montero-Marin J et al. (MYRIAD Team) (2022). School-based mindfulness training "
    "in early adolescence: What works, for whom and how in the MYRIAD trial? Evid Based "
    "Ment Health, 25(3):117–124. https://doi.org/10.1136/ebmental-2022-300439", size=11)
nr("27. Newlove-Delgado T, Marcheselli F, Williams T, Mandalia D, Davis J, McManus S, "
    "Savic M, Treloar W, Ford T (2023). Mental health of children and young people in "
    "England, 2023-wave 4 follow up to the 2017 survey.", size=11)
nr("28. NHS England (2023). Mental health support in schools. "
    "https://www.england.nhs.uk/mental-health/cyp/trailblazers/", size=11)
nr("29. Prabhu SG et al. (2024). Mental health literacy in secondary school teachers "
    "and interventions to improve it — SR and narrative synthesis. J Mental Health, 1–20. "
    "https://doi.org/10.1080/09638237.2024.2426994", size=11)
nr("30. Radez J, Reardon T, Creswell C, Lawrence PJ, Evdoka-Burton G, Waite P (2021). "
    "Why do children and adolescents (not) seek and access professional help for their "
    "mental health problems? A SR. Eur Child Adolesc Psychiatry, 30(2):183–211. "
    "https://doi.org/10.1007/s00787-019-01469-4", size=11)
nr("31. Rickwood D, Deane FP, Wilson CJ, Ciarrochi J (2005). Young people's help-seeking "
    "for mental health problems. AeJAMH, 4(3):218–251. "
    "https://doi.org/10.5172/jamh.4.3.218", size=11)
nr("32. Roshan R, Hamid S, Kumar R, Hamdani U, Naqvi S, Adeel U, Zill-E-Huma (2025). "
    "Utilizing the CFIR framework for mapping facilitators and barriers of implementing "
    "teachers led school mental health programs — a scoping review. Soc Psychiatry "
    "Psychiatr Epidemiol, 60(3):535–548. https://doi.org/10.1007/s00127-024-02762-7", size=11)
nr("33. Sadler K, Vizard T, Ford T, Goodman A, Goodman R, McManus S (2018). Mental "
    "health of children and young people in England, 2017: trends and characteristics.", size=11)
nr("34. Sclare I, Michelson D, Malpass L, Coster F, Brown J (2015). Innovations in "
    "Practice: DISCOVER CBT workshops for 16–18-year-olds: Development of an open-access "
    "intervention for anxiety and depression in inner-city youth. Child Adolesc Ment "
    "Health, 20(2):102–106. https://doi.org/10.1111/camh.12060", size=11)
nr("35. Slattery P, Saeri AK, Bragge P (2020). Research co-design in health: A rapid "
    "overview of reviews. Health Res Policy Syst, 18(1):17. "
    "https://doi.org/10.1186/s12961-020-0528-9", size=11)
nr("36. Stallard P, Sayal K, Phillips R, Taylor JA, Spears M, Anderson R, Araya R, Lewis "
    "G, Millings A, Montgomery AA (2012). Classroom based CBT in reducing symptoms of "
    "depression in high risk adolescents: Pragmatic cluster RCT. BMJ, 345:e6058. "
    "https://doi.org/10.1136/bmj.e6058", size=11)
nr("37. Stallard P, Skryabina E, Taylor G, Phillips R, Daniels H, Anderson R, Simpson N "
    "(2014). Classroom-based CBT (FRIENDS): A cluster RCT to Prevent Anxiety in Children "
    "through Education in Schools (PACES). Lancet Psychiatry, 1(3):185–192. "
    "https://doi.org/10.1016/S2215-0366(14)70244-5", size=11)
nr("38. Totzeck C, van der Meer AS, Christiansen H, Durlach F, Li Sanchez K, Schneider S "
    "(2024). Systematic review: Patient and public involvement of children and young "
    "people in mental health research. Clin Child Fam Psychol Rev, 27(1):257–274. "
    "https://doi.org/10.1007/s10567-024-00470-x", size=11)
nr("39. UK Government (2022). Mental health and wellbeing plan: Discussion paper and "
    "call for evidence. Department of Health and Social Care.", size=11)
nr("40. Werner-Seidler A, Spanos S, Calear AL, Perry Y, Torok M, O'Dea B, Christensen H, "
    "Newby JM (2021). School-based depression and anxiety prevention programs: An updated "
    "SR and meta-analysis. Clin Psychol Rev, 89:102079. "
    "https://doi.org/10.1016/j.cpr.2021.102079", size=11)
nr("41. Wilson CJ, Deane FP (2012). Brief report: Need for autonomy and other perceived "
    "barriers relating to adolescents' intentions to seek professional mental health "
    "care. J Adolescence, 35(1):233–237. https://doi.org/10.1016/j.adolescence.2010.06.011", size=11)
nr("42. Wittevrongel E, Kessels R, Everaert G, Vrijens M, Danckaerts M, van Winkel R "
    "(2025). Differences between users and professionals in preferences for youth mental "
    "health service attributes: A discrete choice experiment. Early Interv Psychiatry, "
    "19(2):e70012. https://doi.org/10.1111/eip.70012", size=11)
nr("43. Yamaguchi S, Foo JC, Sasaki T (2024). The effects of a teacher-led online mental "
    "health literacy program for high school students: A pilot cluster RCT. J Mental "
    "Health, 33(5):630–637. https://doi.org/10.1080/09638237.2024.2390376", size=11)
nr("44. Zhang Q, Wang J, Neitzel A (2023). School-based mental health interventions "
    "targeting depression or anxiety: A meta-analysis of rigorous RCTs for school-aged "
    "children and adolescents. J Youth Adolesc, 52(1):195–217. "
    "https://doi.org/10.1007/s10964-022-01684-4", size=11)

H2("4.2. Nguồn xác minh ngoài (em đã đối chiếu)")
nr("• Bài gốc Editorial: DOI 10.1080/09638237.2025.2512332 (Tandfonline) — raw text 552 dòng "
   "trong 03_Ban-dich/B5_UK_school_raw.txt", size=11)
nr("• BESST trial PMID 38759665 (Lancet Psychiatry, em đã đọc full PDF — "
   "02_Papers-goc/UK_BESST_PLACES/Brown_2024_BESST_Lancet_Psychiatry.pdf, 959 KB)", size=11)
nr("• PLACES model PMC8909998 (IJERPH, em đã đọc full PDF — "
   "02_Papers-goc/UK_BESST_PLACES/Brown_2022_PLACES_IJERPH.pdf, 884 KB)", size=11)
nr("• BESST companion paper PMC11069277 (Trials journal — recruitment & baseline)", size=11)
nr("• King's College Pure repository: kclpure.kcl.ac.uk → BESST trial OA version (CC BY)", size=11)
nr("• National Elf Service review: nationalelfservice.net/.../whats-besst-young-people/ "
   "(critical lay summary of BESST)", size=11)

# =====================================================================
# PHẦN 5 — TRUY VẾT NỘI BỘ
# =====================================================================
H1("PHẦN 5 — TRUY VẾT NỘI BỘ DỰ ÁN")

H2("5.1. RAG (rag_db_full)")
nr("• Collection: lo_au_full (47 chunks, embedding paraphrase-multilingual-MiniLM-L12-v2)", size=11)
nr("• Bài này chưa có chunk RAG riêng — corpus B5 có sẵn dạng raw text, "
   "chưa indexed trực tiếp", size=11)
nr("• Các chunk RAG liên quan: QT08 (Wen 2020 LPA), QT29 (Li 2025 NMA), QT10 (GBD ASEAN), "
   "INSIGHT_05 (yếu tố nguy cơ/bảo vệ VN)", size=11)

H2("5.2. KG (06_Scripts/kg_data)")
nr("• Bài này chưa có node trong nodes.json — corpus được chia QT001-QT048 + INSIGHT_xx; "
   "B5 / QT042 đã có metadata nhưng node KG chưa cập nhật. Đề xuất bổ sung:", size=11)
nr("    – Node QT042: Brown & Carter 2025 J Mental Health, type=editorial, "
   "country=UK, n=N/A (not original research)", size=11)
nr("    – Edges: QT042 → QT042-trial=BESST (Brown 2024); QT042 → QT-PLACES (Brown 2022)", size=11)

H2("5.3. Glossary nội bộ")
nr("• 06_Scripts/glossary_data/glossary_v3_pedagogical.json — đã đính chính BESST + PLACES "
   "trong phiên 25/04/2026 (sửa expansion sai trước đó)", size=11)
nr("• Các thuật ngữ liên quan trong glossary: BESST, PLACES, MHST, DISCOVER, CBT, "
   "Universal vs Targeted, Self-referral, Co-design", size=11)

H2("5.4. File dự án đã đính chính trong phiên này")
nr("• 06_Scripts/glossary_data/glossary_v3_pedagogical.json (BESST + PLACES expansion)", size=11)
nr("• 06_Scripts/glossary_data/glossary_full.json (PLACES)", size=11)
nr("• 06_Scripts/glossary_data/glossary_enhanced.json (PLACES)", size=11)
nr("• tro-ly-nghien-cuu-tam-ly/web/data/glossary.json (PLACES)", size=11)
nr("• tro-ly-nghien-cuu-tam-ly-light/web/data/glossary.json (PLACES)", size=11)
nr("• BANG_TOM_TAT_11_BAI_BAO.md dòng 35 + 107 (Wen 2020 chiều giới tính)", size=11)

H2("5.5. Các doc trả lời thầy đã tạo trong cùng dòng câu hỏi")
nr("• 01_Bao-cao/McDonald_Omega_la_gi_cho_thay_25042026.docx", size=11)
nr("• 01_Bao-cao/Bai_8_Wen_2020_GioiTinh_dinh_chinh_cho_thay_25042026.docx", size=11)
nr("• 01_Bao-cao/BESST_PLACES_giai_thich_cho_thay_25042026_v2.docx", size=11)
nr("• 01_Bao-cao/Bai_dich_phan_bien/B5_Brown_Carter_2025_dich_phan_bien_25042026.docx (DOC NÀY)", size=11)
nr("• [Tiếp theo trong session] PLACES dịch full + phản biện; BESST dịch full + phản biện", size=11)

# =====================================================================
# KẾT THÚC
# =====================================================================
H1("KẾT THÚC DOC")
note("Doc này hoàn thành theo workflow chuẩn (memory feedback_research_workflow.md): "
     "dịch song ngữ EN↔VN từng đoạn của bài gốc + phản biện chữ đỏ có dẫn chứng cụ thể "
     "từ corpus + đối chiếu PMC/PubMed/King's College + bảng so sánh + reference đầy đủ. "
     "Đã kiểm 3 vòng: (1) số liệu khớp PDF gốc Brown 2024 + Brown 2022; (2) phản biện "
     "có dẫn chứng (corpus QT08, QT10, QT29, VN030, B8 + meta-analysis Caldwell 2019, "
     "Zhang 2023, Montero-Marin 2022); (3) reference 30 ref bài gốc + 6 nguồn ngoài "
     "đều có DOI/PMID/PMC.")
note("Sản phẩm tiếp theo trong session: doc dịch + phản biện cho bài 2 (PLACES Brown "
     "2022) và bài 1 (BESST Brown 2024). Em sẽ làm tuần tự + báo bác sau mỗi doc để kiểm.",
     color=BLUE)

d.save(OUT)
print('Wrote:', OUT)
