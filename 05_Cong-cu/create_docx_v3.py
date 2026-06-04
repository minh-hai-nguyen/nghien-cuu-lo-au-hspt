"""
Script v3: Clean docx - Vietnamese text only, real Word tables, no images.
"""
import os, glob
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
TRANS_DIR = os.path.join(BASE, "Translations")
OUT_DIR = os.path.join(BASE, "DocFiles")
os.makedirs(OUT_DIR, exist_ok=True)


def add_word_table(doc, table_lines, in_critical=False):
    data_lines = []
    for tl in table_lines:
        tl = tl.strip()
        if tl and not all(c in '|-: ' for c in tl):
            cells = [c.strip() for c in tl.strip('|').split('|')]
            data_lines.append(cells)
    if not data_lines:
        return
    num_cols = max(len(row) for row in data_lines)
    for row in data_lines:
        while len(row) < num_cols:
            row.append('')

    table = doc.add_table(rows=len(data_lines), cols=num_cols)
    table.style = 'Table Grid'
    # Auto-adjust column widths based on content
    from docx.shared import Cm
    if num_cols > 1:
        # Calculate max content length per column
        col_lens = [0] * num_cols
        for row in data_lines:
            for j, cell in enumerate(row):
                col_lens[j] = max(col_lens[j], len(cell.strip()))
        total = sum(col_lens) or 1
        page_width = 16.0  # cm usable
        for j in range(num_cols):
            w = max(1.5, (col_lens[j] / total) * page_width)
            for row in table.rows:
                row.cells[j].width = Cm(w)
    for i, row_data in enumerate(data_lines):
        for j, cell_text in enumerate(row_data):
            cell = table.cell(i, j)
            cell.text = ''
            p = cell.paragraphs[0]
            clean = cell_text.strip()
            if '**' in clean:
                parts = clean.split('**')
                for k, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(9)
                        run.font.name = 'Times New Roman'
                        if k % 2 == 1:
                            run.bold = True
                        if in_critical:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(clean)
                run.font.size = Pt(9)
                run.font.name = 'Times New Roman'
                if in_critical:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            if i == 0:
                for run in p.runs:
                    run.bold = True
    doc.add_paragraph()


def md_to_docx_v3(md_path, docx_path):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_critical = False
    table_buf = []

    for line in lines:
        line = line.rstrip('\n')

        # Collect table lines
        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            table_buf.append(line)
            continue
        elif table_buf:
            add_word_table(doc, table_buf, in_critical)
            table_buf = []

        if not line.strip():
            continue

        # Handle image insertion markers
        if line.strip().startswith('<!-- INSERT_IMAGE:'):
            try:
                parts = line.strip().replace('<!-- INSERT_IMAGE:', '').replace('-->', '').strip().split('|')
                img_rel = parts[0].strip()
                caption = parts[1].strip() if len(parts) > 1 else ''
                img_path = os.path.join(BASE, img_rel)
                if os.path.exists(img_path):
                    from docx.shared import Inches
                    doc.add_picture(img_path, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = 1  # center
                    if caption:
                        p = doc.add_paragraph()
                        run = p.add_run(caption)
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
                        p.alignment = 1
                    doc.add_paragraph()
            except Exception as e:
                doc.add_paragraph(f'[Image error: {e}]')
            continue

        if 'QUAN DIEM PHAN BIEN' in line.upper() or 'CRITICAL REVIEW' in line:
            in_critical = True

        # Page markers (orange)
        if line.strip().startswith('**--- Trang') and line.strip().endswith('---**'):
            p = doc.add_paragraph()
            marker_text = line.strip().strip('*').strip()
            run = p.add_run(marker_text)
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
            continue


        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            if in_critical:
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('---'):
            pass
        elif line.lstrip().startswith('- '):
            indent = len(line) - len(line.lstrip())
            text = line.lstrip(' -')
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                parts = text.split('**')
                for k, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                        if k % 2 == 1:
                            run.bold = True
                        if in_critical:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                if in_critical:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        else:
            p = doc.add_paragraph()
            text = line
            if '**' in text:
                parts = text.split('**')
                for k, part in enumerate(parts):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(12)
                        if k % 2 == 1:
                            run.bold = True
                        if in_critical:
                            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            else:
                run = p.add_run(text)
                run.font.size = Pt(12)
                if in_critical:
                    run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)

    if table_buf:
        add_word_table(doc, table_buf, in_critical)

    doc.save(docx_path)


# Known abbreviations to highlight red on first appearance
KNOWN_ABBREVS = [
    'PHQ-9A', 'PHQ-9', 'GAD-10', 'GAD-7', 'DASS-21', 'CES-D', 'SAS',
    'EMBU', 'LPA', 'SRSMSS', 'PSQI', 'IGDS9-SF', 'IGD', 'MPVS',
    'AOR', 'SPSS', 'WHO', 'DSM-IV', 'SPIN', 'ASSIST', 'ANOVA',
    'NCD', 'MHT', 'BIC', 'AIC', 'BLRT', 'COVID-19', 'PPS', 'IRB',
    'KSA', 'SEAR', 'GMCH', 'PTSD', 'COR', 'OR', 'CI',
]


def render_text_with_abbrev_highlight(p, text, font_size, in_critical, seen_abbrevs):
    """Render text, highlighting abbreviations red on first appearance."""
    import re
    # Find abbreviation patterns: ABBREV (explanation)
    # Split text around known abbreviations
    remaining = text
    while remaining:
        # Find earliest abbreviation match
        earliest_pos = len(remaining)
        earliest_abbr = None
        for abbr in KNOWN_ABBREVS:
            pos = remaining.find(abbr)
            if pos != -1 and pos < earliest_pos:
                # Make sure it's a whole word (not part of larger word)
                end_pos = pos + len(abbr)
                if end_pos < len(remaining) and remaining[end_pos].isalpha():
                    continue
                earliest_pos = pos
                earliest_abbr = abbr

        if earliest_abbr is None:
            # No more abbreviations, render rest as normal
            if remaining:
                run = p.add_run(remaining)
                run.font.size = font_size
                if in_critical:
                    run.font.color.rgb = RGBColor(0xCC, 0, 0)
            break

        # Text before abbreviation
        before = remaining[:earliest_pos]
        if before:
            run = p.add_run(before)
            run.font.size = font_size
            if in_critical:
                run.font.color.rgb = RGBColor(0xCC, 0, 0)

        # The abbreviation itself
        is_first = earliest_abbr not in seen_abbrevs
        run = p.add_run(earliest_abbr)
        run.font.size = font_size
        if is_first:
            run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)  # Red
            run.bold = True
            seen_abbrevs.add(earliest_abbr)
        elif in_critical:
            run.font.color.rgb = RGBColor(0xCC, 0, 0)

        remaining = remaining[earliest_pos + len(earliest_abbr):]


def md_to_docx_full(md_path, docx_path):
    """Enhanced version: adds QR, link, abbreviations."""
    basename = os.path.splitext(os.path.basename(md_path))[0]

    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)

    # Add header with link + QR code
    add_header_with_qr(doc, basename)

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    in_critical = False
    table_buf = []
    found_critical = False
    seen_abbrevs = set()  # Track first appearance of abbreviations

    for line in lines:
        line = line.rstrip('\n')

        if line.strip().startswith('|') and '|' in line.strip()[1:]:
            table_buf.append(line)
            continue
        elif table_buf:
            add_word_table(doc, table_buf, in_critical)
            table_buf = []

        if not line.strip():
            continue

        # Handle image markers
        if line.strip().startswith('<!-- INSERT_IMAGE:'):
            try:
                parts = line.strip().replace('<!-- INSERT_IMAGE:', '').replace('-->', '').strip().split('|')
                img_rel = parts[0].strip()
                caption = parts[1].strip() if len(parts) > 1 else ''
                img_path = os.path.join(BASE, img_rel)
                if os.path.exists(img_path):
                    from docx.shared import Inches
                    doc.add_picture(img_path, width=Inches(5.5))
                    doc.paragraphs[-1].alignment = 1
                    if caption:
                        p = doc.add_paragraph()
                        run = p.add_run(caption)
                        run.font.size = Pt(9)
                        run.font.italic = True
                        run.font.color.rgb = RGBColor(0x33, 0x33, 0x99)
                        p.alignment = 1
                    doc.add_paragraph()
            except:
                pass
            continue

        # Insert abbreviations BEFORE critical review
        if ('QUAN DIEM PHAN BIEN' in line.upper() or 'CRITICAL REVIEW' in line) and not found_critical:
            add_abbreviations_footnote(doc, basename)
            found_critical = True
            in_critical = True

        # Page markers (orange)
        if line.strip().startswith('**--- Trang') and line.strip().endswith('---**'):
            p = doc.add_paragraph()
            marker_text = line.strip().strip('*').strip()
            run = p.add_run(marker_text)
            run.font.size = Pt(10)
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)
            continue


        # Headings
        if line.startswith('# '):
            doc.add_heading(line[2:].strip(), level=1)
        elif line.startswith('## '):
            h = doc.add_heading(line[3:].strip(), level=2)
            if in_critical:
                for r in h.runs:
                    r.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        elif line.startswith('### '):
            doc.add_heading(line[4:].strip(), level=3)
        elif line.startswith('#### '):
            doc.add_heading(line[5:].strip(), level=4)
        elif line.startswith('---'):
            pass
        elif line.lstrip().startswith('- '):
            text = line.lstrip(' -')
            p = doc.add_paragraph(style='List Bullet')
            if '**' in text:
                for k, part in enumerate(text.split('**')):
                    if part:
                        run = p.add_run(part)
                        run.font.size = Pt(11)
                        if k % 2 == 1: run.bold = True
                        if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
            else:
                run = p.add_run(text)
                run.font.size = Pt(11)
                if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
        else:
            p = doc.add_paragraph()
            if '**' in line:
                for k, part in enumerate(line.split('**')):
                    if part:
                        if k % 2 == 1:
                            run = p.add_run(part)
                            run.font.size = Pt(12)
                            run.bold = True
                            if in_critical: run.font.color.rgb = RGBColor(0xCC, 0, 0)
                        else:
                            render_text_with_abbrev_highlight(p, part, Pt(12), in_critical, seen_abbrevs)
            else:
                render_text_with_abbrev_highlight(p, line, Pt(12), in_critical, seen_abbrevs)

    if table_buf:
        add_word_table(doc, table_buf, in_critical)

    doc.save(docx_path)


# Paper metadata: DOI links, QR codes, abbreviations
PAPER_META = {
    "01_Jenkins_et_al_2023": {
        "url": "https://doi.org/10.1177/00207640221140282",
        "qr": "QR_Codes/01_qr.png",
        "abbrevs": {
            "PHQ-9A": "Patient Health Questionnaire for Adolescents (Bang cau hoi Suc khoe Benh nhan danh cho Thanh thieu nien)",
            "GAD-10": "Generalized Anxiety Disorder 10-item scale (Thang do Roi loan Lo au Tong quat 10 muc)",
            "IQR": "Interquartile Range (Khoang tu phan vi)",
            "IRB": "Institutional Review Board (Hoi dong Xet duyet Nghien cuu)",
            "UCSD": "University of California San Diego",
        }
    },
    "02_Saikia_et_al_2023": {
        "url": "https://doi.org/10.4103/ijcm.ijcm_614_21",
        "qr": "QR_Codes/02_qr.png",
        "abbrevs": {
            "DASS-21": "Depression Anxiety Stress Scale-21 Items (Thang do Tram cam Lo au Cang thang 21 muc)",
            "WHO": "World Health Organization (To chuc Y te The gioi)",
            "SEAR": "South-East Asia Region (Khu vuc Dong Nam A)",
            "NCD": "Non-Communicable Diseases (Benh khong lay nhiem)",
            "SPSS": "Statistical Package for Social Sciences",
            "GMCH": "Gauhati Medical College and Hospital",
        }
    },
    "03_Mandaknalli_Malusare_2021": {
        "url": "https://www.medpulse.in/Psychology/Article/Volume18Issue3/Psy_18_3_2.pdf",
        "qr": "QR_Codes/03_qr.png",
        "abbrevs": {
            "GAD-7": "Generalized Anxiety Disorder 7-item scale (Thang do Roi loan Lo au Tong quat 7 muc)",
            "SPSS": "Statistical Package for Social Sciences",
            "PTSD": "Post-Traumatic Stress Disorder (Roi loan cang thang sau chan thuong)",
        }
    },
    "05_Alharbi_et_al_2019": {
        "url": "https://doi.org/10.4103/jfmpc.jfmpc_383_18",
        "qr": "QR_Codes/05_qr.png",
        "abbrevs": {
            "PHQ-9": "Patient Health Questionnaire-9 (Bang cau hoi Suc khoe Benh nhan 9 muc)",
            "GAD-7": "Generalized Anxiety Disorder 7-item scale (Thang do Roi loan Lo au Tong quat 7 muc)",
            "DSM-IV": "Diagnostic and Statistical Manual of Mental Disorders, 4th Edition",
            "DASS": "Depression Anxiety Stress Scale (Thang do Tram cam Lo au Cang thang)",
            "SPSS": "Statistical Package for Social Sciences",
            "KSA": "Kingdom of Saudi Arabia (Vuong quoc A Rap Saudi)",
        }
    },
    "06_Nakie_et_al_2022": {
        "url": "https://doi.org/10.1186/s12888-022-04393-1",
        "qr": "QR_Codes/06_qr.png",
        "abbrevs": {
            "DASS-21": "Depression Anxiety Stress Scale-21 Items",
            "SPIN": "Social Phobia Inventory (Bang kiem Am anh Xa hoi)",
            "ASSIST": "Alcohol, Smoking, Substance Involvement Screening Test",
            "AOR": "Adjusted Odds Ratio (Ty so chenh da hieu chinh)",
            "CI": "Confidence Interval (Khoang tin cay)",
            "COR": "Crude Odds Ratio (Ty so chenh tho)",
        }
    },
    "07_Chen_et_al_2023": {
        "url": "https://doi.org/10.1186/s12888-023-05068-1",
        "qr": "QR_Codes/07_qr.png",
        "abbrevs": {
            "PHQ-9": "Patient Health Questionnaire-9",
            "GAD-7": "Generalized Anxiety Disorder 7-item scale",
            "MPVS": "Multidimensional Peer-Victimization Scale (Thang do Da chieu ve Nan nhan Bat nat)",
            "PSQI": "Pittsburgh Sleep Quality Index (Chi so Chat luong Giac ngu Pittsburgh)",
            "IGDS9-SF": "Internet Gaming Disorder Scale 9-Short Form",
            "IGD": "Internet Gaming Disorder (Roi loan Choi game Truc tuyen)",
            "OR": "Odds Ratio (Ty so chenh)",
        }
    },
    "08_Wen_et_al_2020": {
        "url": "https://doi.org/10.3390/ijerph17114079",
        "qr": "QR_Codes/08_qr.png",
        "abbrevs": {
            "LPA": "Latent Profile Analysis (Phan tich Ho so Tiem an)",
            "MHT": "Mental Health Test (Trac nghiem Suc khoe Tam than)",
            "AIC": "Akaike Information Criterion",
            "BIC": "Bayesian Information Criterion",
            "OR": "Odds Ratio (Ty so chenh)",
        }
    },
    "09_Qiu_et_al_2022": {
        "url": "https://doi.org/10.3389/fpsyg.2022.897339",
        "qr": "QR_Codes/09_qr.png",
        "abbrevs": {
            "LPA": "Latent Profile Analysis (Phan tich Ho so Tiem an)",
            "EMBU": "Egna Minnen Betraffande Uppfostran (Thang do Phong cach Nuoi day con)",
            "CES-D": "Center for Epidemiologic Studies Depression Scale",
            "SAS": "Self-Rating Anxiety Scale (Thang do Lo au Tu danh gia)",
            "SRSMSS": "Self-rating Resilience Scale for Middle School Students",
            "OR": "Odds Ratio (Ty so chenh)",
            "BLRT": "Bootstrapped Likelihood Ratio Test",
        }
    },
    "10_Xu_et_al_2021": {
        "url": "https://doi.org/10.1016/j.jad.2021.03.080",
        "qr": "QR_Codes/10_qr.png",
        "abbrevs": {
            "GAD-7": "Generalized Anxiety Disorder 7-item scale",
            "COVID-19": "Coronavirus Disease 2019",
            "OR": "Odds Ratio (Ty so chenh)",
            "CI": "Confidence Interval (Khoang tin cay)",
            "SPSS": "Statistical Package for Social Sciences",
        }
    },
    "11_Bhardwaj_et_al_2020": {
        "url": "https://gmch.gov.in/sites/default/files/documents/DAS.pdf",
        "qr": "QR_Codes/11_qr.png",
        "abbrevs": {
            "DASS-21": "Depression Anxiety Stress Scale-21 Items",
            "PPS": "Probability Proportional to Size (Xac suat ty le voi co mau)",
            "ANOVA": "Analysis of Variance (Phan tich Phuong sai)",
            "WHO": "World Health Organization (To chuc Y te The gioi)",
        }
    },
}


def add_header_with_qr(doc, basename):
    """Add link, QR code, and abbreviations footnote at start of doc."""
    meta = PAPER_META.get(basename)
    if not meta:
        return

    # Check if online journal
    online_journals = {"06_", "07_", "08_", "09_"}
    is_online = any(basename.startswith(prefix) for prefix in online_journals)

    # Add link
    p = doc.add_paragraph()
    run = p.add_run("Link bai bao goc: ")
    run.font.size = Pt(10)
    run.bold = True
    run = p.add_run(meta["url"])
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(0x00, 0x00, 0xCC)
    run.underline = True

    if is_online:
        p2 = doc.add_paragraph()
        run2 = p2.add_run("* Day la tap chi truc tuyen (online/open access). Bai bao duoc xuat ban dang dien tu, khong co so trang tap chi truyen thong. So trang trong ban dich nay la so thu tu trang cua bai bao (Page X of Y).")
        run2.font.size = Pt(9)
        run2.font.italic = True
        run2.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    # Add QR code
    qr_path = os.path.join(BASE, meta["qr"])
    if os.path.exists(qr_path):
        try:
            from docx.shared import Inches
            doc.add_picture(qr_path, width=Inches(1.2))
            doc.paragraphs[-1].alignment = 0  # left align
            p = doc.add_paragraph()
            run = p.add_run("Quet QR de truy cap bai bao goc")
            run.font.size = Pt(8)
            run.font.italic = True
        except:
            pass

    doc.add_paragraph()  # spacing


def add_abbreviations_footnote(doc, basename):
    """Add abbreviations table at the end (before critical review)."""
    meta = PAPER_META.get(basename)
    if not meta or not meta.get("abbrevs"):
        return

    doc.add_paragraph()
    h = doc.add_heading("CHU THICH TU VIET TAT / ABBREVIATIONS", level=3)

    abbrevs = meta["abbrevs"]
    table = doc.add_table(rows=len(abbrevs) + 1, cols=2)
    table.style = 'Table Grid'

    # Header
    table.cell(0, 0).text = "Viet tat"
    table.cell(0, 1).text = "Giai thich"
    for run in table.cell(0, 0).paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)
    for run in table.cell(0, 1).paragraphs[0].runs:
        run.bold = True
        run.font.size = Pt(9)

    for i, (abbr, meaning) in enumerate(abbrevs.items(), 1):
        cell0 = table.cell(i, 0)
        cell1 = table.cell(i, 1)
        cell0.text = ""
        cell1.text = ""
        run0 = cell0.paragraphs[0].add_run(abbr)
        run0.font.size = Pt(9)
        run0.bold = True
        run1 = cell1.paragraphs[0].add_run(meaning)
        run1.font.size = Pt(9)

    # Set column widths
    from docx.shared import Cm
    for row in table.rows:
        row.cells[0].width = Cm(3)
        row.cells[1].width = Cm(13)

    doc.add_paragraph()


# Process all files
print("Tao DOCX v3 - chi text tieng Viet + bang Word that")
for md_file in sorted(glob.glob(os.path.join(TRANS_DIR, "*.md"))):
    name = os.path.splitext(os.path.basename(md_file))[0]
    out = os.path.join(OUT_DIR, f"{name}.docx")
    if name in PAPER_META:
        md_to_docx_full(md_file, out)
    else:
        md_to_docx_v3(md_file, out)
    print(f"  OK: {name}.docx")

# Summary
md_to_docx_v3(os.path.join(BASE, "BANG_TOM_TAT_11_BAI_BAO.md"),
              os.path.join(OUT_DIR, "BANG_TOM_TAT_TONG_HOP.docx"))
print("  OK: BANG_TOM_TAT_TONG_HOP.docx")
print("HOAN TAT!")
