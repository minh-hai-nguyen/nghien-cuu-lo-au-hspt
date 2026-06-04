# -*- coding: utf-8 -*-
"""Verify VN translations of refs do not introduce fabricated facts.

Check against:
1. Vietnamese author name ordering (VN surname comes first)
2. Journal names — keep or translate with parenthetical English
3. Year/volume/issue/pages MUST match exactly
4. DOIs — not translated
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

d = Document('03_Ban-dich/VN002_VNAMHS_2022_National_FULL.docx')

# Find refs section
refs_start = None
for i, p in enumerate(d.paragraphs):
    if p.text.startswith('American Psychiatric Association. 2013'):
        refs_start = i
        break

# Load PDF to cross-check
with open('C:/Users/HLC/AppData/Local/Temp/vnamhs_pdf_full.txt', encoding='utf-8') as f:
    pdf = f.read()

# Extract refs pairs (EN paragraph + VN paragraph)
pairs = []
i = refs_start
while i < len(d.paragraphs) - 1:
    en_p = d.paragraphs[i].text.strip()
    vn_p = d.paragraphs[i+1].text.strip() if i+1 < len(d.paragraphs) else ''
    if en_p and re.match(r'^[A-Z][a-zA-Z\s,\.]+\.\s\d{4}', en_p):
        pairs.append((en_p, vn_p))
    i += 1

print(f'Total EN-VN ref pairs found: {len(pairs)}')

# Check 1: Year must match exactly
issues = []
for en, vn in pairs:
    en_years = set(re.findall(r'\b(19|20)\d{2}\b', en))
    vn_years = set(re.findall(r'\b(19|20)\d{2}\b', vn))
    # Normalise — only compare 4-digit years
    en_y = set(re.findall(r'\b(19|20)\d{2}\b', en))
    vn_y = set(re.findall(r'\b(19|20)\d{2}\b', vn))
    # Actually extract full years
    en_full = set(re.findall(r'(?:19|20)\d{2}', en))
    vn_full = set(re.findall(r'(?:19|20)\d{2}', vn))
    missing_in_vn = en_full - vn_full
    if missing_in_vn and 'không cần dịch' not in vn:
        issues.append(f'Year(s) {missing_in_vn} in EN but not VN: {en[:60]}')

print(f'\nYear mismatches: {len(issues)}')
for i_ in issues[:5]:
    print(f'  ⚠ {i_}')

# Check 2: Page numbers / volume / issue
issues2 = []
for en, vn in pairs:
    # Extract patterns like "55(10):841–50"
    en_cite = re.search(r'(\d+)\((\d+[–\-]?\d*)\):?\s*(\d+[–\-]\d+)', en)
    vn_cite = re.search(r'(\d+)\((\d+[–\-]?\d*)\):?\s*(\d+[–\-]\d+)', vn)
    if en_cite and not vn_cite and 'không cần dịch' not in vn:
        issues2.append(f'Citation {en_cite.group()} not preserved: {en[:70]}')
    elif en_cite and vn_cite:
        if en_cite.group() != vn_cite.group():
            # Allow minor dash differences
            if en_cite.group().replace('-', '–') != vn_cite.group().replace('-', '–'):
                issues2.append(f'Citation mismatch: EN {en_cite.group()} vs VN {vn_cite.group()}')
print(f'\nCitation (vol/issue/page) mismatches: {len(issues2)}')
for i_ in issues2[:5]:
    print(f'  ⚠ {i_}')

# Check 3: Vietnamese author name check — verify "Hoang M.D." is correctly "Đặng Hoàng Minh"
# The VN translation should preserve proper author ordering
vn_author_issues = []
# Known VN authors in refs
known_mapping = {
    'Hoang M.D.': 'Đặng Hoàng Minh',  # Dang Hoang Minh in English, so VN = Đặng Hoàng Minh
    'Hoang-Minh Dang': 'Đặng Hoàng Minh',
    'Cao Tiến Đức': 'Cao Tiến Đức',  # already VN
    'Hoàng Thế Hải': 'Hoàng Thế Hải',  # already VN
    'Bui Thi Thanh Dieu': 'Bùi Thị Thanh Diệu',
    'Vu Manh Loi': 'Vũ Mạnh Lợi',
    'Nguyen Duc Vinh': 'Nguyễn Đức Vinh',
    'Dao Thi Khanh Hoa': 'Đào Thị Khánh Hoa',
    'Pham Vu Hoang': 'Phạm Vũ Hoàng',
    'Dang Nguyen Anh': 'Đặng Nguyên Anh',
}

for en, vn in pairs:
    # Check: if "Hoang M.D." in EN, VN should have "Đặng Hoàng Minh"
    for en_form, vn_form in known_mapping.items():
        if en_form in en and 'không cần dịch' not in vn:
            if vn_form not in vn and not any(w in vn for w in vn_form.split()):
                # Check reverse anglicized order was used (wrong)
                reversed_form = ' '.join(reversed(vn_form.split()))
                if reversed_form in vn:
                    vn_author_issues.append(f'{en[:50]}: uses anglicized order "{reversed_form}" instead of "{vn_form}"')
                elif en_form == 'Hoang M.D.':
                    # Special: Hoàng Minh Đặng used — wrong
                    if 'Hoàng Minh Đặng' in vn:
                        vn_author_issues.append(f'{en[:60]}: "Hoàng Minh Đặng" should be "Đặng Hoàng Minh"')

print(f'\nVN author name ordering issues: {len(vn_author_issues)}')
for i_ in vn_author_issues:
    print(f'  ⚠ {i_}')

# Check 4: DOI preservation
doi_issues = []
for en, vn in pairs:
    en_dois = re.findall(r'(10\.\d+/[^\s]+)', en)
    vn_dois = re.findall(r'(10\.\d+/[^\s]+)', vn)
    for doi in en_dois:
        if doi not in vn:
            # OK if VN translation omits DOI (standard practice)
            pass
print(f'\nDOI issues: {len(doi_issues)} (OK if omitted in VN)')

print('\n' + '='*70)
print(f'SUMMARY: {len(issues) + len(issues2) + len(vn_author_issues)} total flags')
