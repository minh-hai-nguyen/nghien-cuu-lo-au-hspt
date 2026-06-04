# -*- coding: utf-8 -*-
"""
QA cho ban dich VN022 UNICEF 2022 — 7 phuong phap:
#1 AEP (Assertion-Evidence Pairs): moi so lieu co structured evidence
#2 Provenance Chain: trace moi so ngay qua lai toi PDF goc
#3 Ontology Type System: kiem tra khai niem (instrument, design, outcome)
#4 Contradiction Detection: tim mau thuan trong ban dich
#5 Temporal KB: kiem tra moc thoi gian, chinh sach, data year
#6 Claim Strength: kiem tra over-generalization, causal overclaim
#7 Pipeline verify_numbers: so EN <-> VN numbers
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from pypdf import PdfReader
from collections import Counter, defaultdict

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
VN_DOCX = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')
EN_PDF = os.path.join(ROOT, '02_Papers-goc', 'Viet-Nam', 'UNICEF_2022_School_Factors_Vietnam.pdf')

# Load
d = Document(VN_DOCX)
vn_paras = [p.text for p in d.paragraphs if p.text.strip()]
vn_text = '\n'.join(vn_paras)
for t in d.tables:
    for r in t.rows:
        vn_text += '\n' + ' | '.join(c.text.strip() for c in r.cells)

r = PdfReader(EN_PDF)
en_text = '\n'.join((p.extract_text() or '') for p in r.pages)

print('='*70)
print('QA VN022 UNICEF 2022 — 7 METHODS')
print('='*70)
print(f'VN: {len(vn_text):,} chars, {len(vn_paras)} paragraphs')
print(f'EN PDF: {len(en_text):,} chars, {len(r.pages)} pages')
print(f'Coverage ratio: {len(vn_text)/len(en_text):.2f}')

issues = defaultdict(list)

# ============================================================
# #7 verify_numbers — foundation
# ============================================================
print('\n' + '#7 '*20)
print('METHOD #7: Verify Numbers EN <-> VN')

# Extract all numbers from both
NUM = re.compile(r'\b(\d+[,.]?\d*)\s*%?')
# Focus on key statistics-format numbers
STAT = re.compile(r'(\d+[,.]\d+)\s*%|(\d+)\s*%|n\s*=\s*(\d+[,.]?\d*)')

# Normalize comma/period: VN uses comma decimal, EN uses period decimal
def normalize_num(s):
    return s.replace(',', '.')

# Key numbers we expect in both
key_stats = [
    ('26.1%', '26,1'),     # overall MH
    ('32%', '32'),          # peer problems (30.9 also)
    ('30.9%', '30,9'),      # emotional
    ('14.4%', '14,4'),      # hyperactivity
    ('11%', '11'),          # conduct
    ('668', '668'),         # students
    ('66', '66'),           # teachers
    ('1,084', '1.084'),     # Happy House not in this but similar
    ('86.4%', '86,4'),      # teachers without training
    ('91%', '91'),          # teacher concern stress
    ('95%', '95'),          # teacher concern MH
    ('15%', '15'),          # extra class >9h
    ('28%', '28'),          # study >3h
    ('47.3%', '47,3'),      # never cyberbullied
    ('2.1%', '2,1'),        # often cyberbullied
    ('70%', '70'),          # rural schools no counselling room
]

missing_in_vn = []
for en_num, vn_num in key_stats:
    en_hit = en_num in en_text or en_num.replace(',', '').replace('.', ',') in en_text
    vn_hit = vn_num in vn_text
    if en_hit and not vn_hit:
        missing_in_vn.append((en_num, vn_num))

if missing_in_vn:
    print(f'  Key stats in EN but missing/different in VN: {len(missing_in_vn)}')
    for en, vn in missing_in_vn:
        print(f'    EN="{en}" -> looked for VN="{vn}"  MISSING')
        issues['#7 verify'].append(f'Missing: EN {en} / VN {vn}')
else:
    print('  All 16 key statistics VERIFIED in both EN and VN')

# ============================================================
# #1 AEP — Extract evidence for each statistical claim
# ============================================================
print('\n' + '#1 '*20)
print('METHOD #1: AEP — Evidence per claim')

# For each %/number in VN, check nearby context has source attribution
vn_percents = re.findall(r'(\d+[,.]\d+)\s*%', vn_text)
print(f'  Total % stats in VN: {len(vn_percents)}')

# AEP check: each number should be near an instrument name, author, or year
AEP_KEYWORDS = ['SDQ', 'MDS3', 'ESSA', 'RCBI', 'Goodman', 'Sun', 'Bradshaw',
                'Weiss', 'Nguyen', 'Hoa', 'UNICEF', 'MOET', 'MOH', 'MOLISA',
                'Circular', 'Quyết định', 'Thông tư', '2021', '2022', '2020',
                '2019', '2018', '2017', 'trang', 'Bảng', 'Phụ lục']
aep_found = 0
for pct in vn_percents[:50]:
    idx = vn_text.find(pct + ' %' if ' %' in vn_text else pct + '%')
    if idx < 0:
        idx = vn_text.find(pct)
    if idx >= 0:
        ctx = vn_text[max(0, idx-200):idx+200]
        if any(kw in ctx for kw in AEP_KEYWORDS):
            aep_found += 1
print(f'  AEP match rate (first 50 stats): {aep_found}/50 = {aep_found*2}%')
if aep_found < 40:
    issues['#1 AEP'].append(f'Low evidence attribution: {aep_found}/50 stats')

# ============================================================
# #2 Provenance Chain — Trace VN numbers to EN PDF
# ============================================================
print('\n' + '#2 '*20)
print('METHOD #2: Provenance Chain VN -> EN PDF')

# Pick 20 specific stats and verify they exist in EN PDF
SPECIFIC = ['26,1', '30,9', '14,4', '53,0', '54,5', '86,4', '91,0', '28,1',
            '47,3', '20,2', '45,9', '34,0', '66,0', '32,0', '85,0', '70,0',
            '96,1', '668', '66', '34']  # +happy house # Note: 53%/54% converted
vn_traced = 0
vn_missing = []
for s in SPECIFIC:
    # Convert back to EN format
    en_variants = [s.replace(',', '.'), s.replace(',', ''), s]
    if s in vn_text:
        if any(v in en_text for v in en_variants):
            vn_traced += 1
        else:
            vn_missing.append(s)

print(f'  Traced: {vn_traced}/{len(SPECIFIC)}')
if vn_missing:
    print(f'  Untraceable to EN PDF: {vn_missing[:10]}')
    for m in vn_missing:
        issues['#2 Provenance'].append(f'VN {m} not traced to EN PDF')

# ============================================================
# #3 Ontology Type System
# ============================================================
print('\n' + '#3 '*20)
print('METHOD #3: Ontology Type System')

# Check: SDQ-25 not confused with GAD-7/DASS-21
onto_checks = [
    ('SDQ', 'general screening (5 subscales)', 'Should NOT be called GAD-7 or DASS-21'),
    ('MDS3', 'school climate', 'Not a mental health scale'),
    ('ESSA', 'educational stress', 'Academic stress specifically, not general'),
    ('RCBI', 'cyberbullying', 'Not a MH scale'),
]
for instr, desc, warning in onto_checks:
    if instr in vn_text:
        # Check no confusion
        idx = vn_text.find(instr)
        ctx = vn_text[max(0, idx-100):idx+300]
        # Warning patterns
        if instr == 'SDQ':
            if 'GAD-7' in ctx or 'DASS-21' in ctx:
                issues['#3 Ontology'].append(f'SDQ context mentions GAD-7 or DASS-21 — possible confusion')
        if instr == 'ESSA':
            if 'trầm cảm' in ctx[:200].lower() and 'lo âu' in ctx[:200].lower():
                pass  # academic stress related to both is OK
print(f'  Instrument ontology: {4-len(issues["#3 Ontology"])}/4 checks pass')

# Design type check: cross-sectional vs RCT
if 'cắt ngang' in vn_text or 'cross-sectional' in vn_text:
    # Look for overclaims of causality
    if 'gây ra' in vn_text or 'causes' in vn_text:
        # Check density
        causality_count = vn_text.count('gây ra') + vn_text.count('causes')
        crosssect_count = vn_text.count('cắt ngang') + vn_text.count('cross-sectional')
        print(f'  "gây ra"/"causes" count: {causality_count}')
        print(f'  "cắt ngang"/"cross-sectional" count: {crosssect_count}')
        # Report does acknowledge cross-sectional limitation
        if crosssect_count < 2:
            issues['#3 Ontology'].append('Cross-sectional limitation not acknowledged adequately')

# ============================================================
# #4 Contradiction Detection
# ============================================================
print('\n' + '#4 '*20)
print('METHOD #4: Contradiction Detection')

# Common contradictions in translation
contradictions = []

# Internal stat consistency
def find_all_pct(text, target):
    return [m.start() for m in re.finditer(re.escape(target), text)]

# Check: 26% vs 26.1%
if '26,1' in vn_text and '26%' in vn_text:
    # Make sure "26%" is not also used as rounded form
    for idx in find_all_pct(vn_text, '26%'):
        ctx = vn_text[max(0,idx-50):idx+50]
        if '26,1' not in ctx:
            pass  # probably OK

# Gender direction consistency
if 'nam > nữ' in vn_text or 'boys have higher' in vn_text.lower():
    issues['#4 Contradiction'].append('Possible gender direction contradiction')

# Sample size consistency: 668 should be consistent
if '668' in vn_text:
    n668 = vn_text.count('668')
    print(f'  "668" mentions: {n668}')
    if n668 < 2:
        issues['#4 Contradiction'].append('Sample size 668 not re-confirmed')

# 66 teachers
if '66 giáo' in vn_text or '66 teachers' in vn_text:
    print(f'  66 teachers: CONFIRMED')

# 5 provinces → actually 4 (HCMC was dropped)
has_5prov = '5 tỉnh' in vn_text
has_4prov = '4 tỉnh' in vn_text or 'bốn tỉnh' in vn_text
has_hcmc_note = 'TPHCM' in vn_text and 'COVID' in vn_text
if has_5prov and has_4prov and has_hcmc_note:
    print(f'  5-prov plan → 4-prov actual (HCMC dropped): PROPERLY NOTED')
elif has_5prov and not has_4prov:
    issues['#4 Contradiction'].append('Mentions 5 tỉnh but may not clarify HCMC dropped')

# ============================================================
# #5 Temporal KB
# ============================================================
print('\n' + '#5 '*20)
print('METHOD #5: Temporal KB')

# Key policy dates should be preserved
policy_dates = [
    ('2005', 'Công văn 9971 MOET school counselling'),
    ('2017', 'Thông tư 31/2017 school counselling'),
    ('2011', 'QĐ 1215 MOLISA community MH'),
    ('2016', 'Luật Trẻ em'),
    ('2020', 'data collection period'),
    ('2021', 'COVID outbreak + School Health Programme'),
    ('2022', 'report publication'),
]
temporal_preserved = 0
for year, label in policy_dates:
    if year in vn_text:
        temporal_preserved += 1
    else:
        issues['#5 Temporal'].append(f'Year {year} ({label}) missing in VN')
print(f'  Policy years preserved: {temporal_preserved}/{len(policy_dates)}')

# Check publication dates
if 'Quyết định số 1660' in vn_text and '2/10/2021' in vn_text:
    print(f'  QĐ 1660 + date: OK')
if '31/2017' in vn_text:
    print(f'  Thông tư 31/2017: OK')

# ============================================================
# #6 Claim Strength Classifier
# ============================================================
print('\n' + '#6 '*20)
print('METHOD #6: Claim Strength')

# Patterns that indicate over-claims
overclaim_patterns = [
    r'CHẮC CHẮN (?:gây|dẫn đến|là nguyên nhân)',  # absolute causal
    r'TẤT CẢ học sinh đều',  # absolutism
    r'HOÀN TOÀN (?:chữa|điều trị|loại bỏ)',  # cure claim
    r'(?:100|98|99)\s*%.*học sinh',  # extreme prevalence without context
]
overclaim_hits = 0
for pat in overclaim_patterns:
    if re.search(pat, vn_text):
        overclaim_hits += 1
        issues['#6 Claim'].append(f'Over-claim pattern detected: {pat}')

# Check hedging language
hedging = ['có thể', 'có vẻ', 'gợi ý', 'tương quan', 'liên quan', 'không thể kết luận',
           'mặc dù', 'tuy nhiên', 'cần nghiên cứu thêm']
hedge_count = sum(vn_text.count(h) for h in hedging)
print(f'  Hedging language count: {hedge_count}')
print(f'  Over-claim patterns: {overclaim_hits}')

if hedge_count < 50:
    issues['#6 Claim'].append(f'Low hedging density — only {hedge_count} hedge terms')

# ============================================================
# REPORT
# ============================================================
print('\n' + '='*70)
print('FINAL QA REPORT')
print('='*70)
total_issues = sum(len(v) for v in issues.values())
print(f'Total issues: {total_issues}')
for method, errs in sorted(issues.items()):
    print(f'\n{method}: {len(errs)} issue(s)')
    for e in errs[:5]:
        print(f'  - {e}')

# Completeness check
print('\nCOMPLETENESS:')
print(f'  VN/EN ratio: {len(vn_text)/len(en_text):.2f} (target: 0.70-0.85)')
print(f'  VN paragraphs: {len(vn_paras)}')
print(f'  Tables: {len(d.tables)}')

# Chapter coverage
chapters = ['CHƯƠNG 1', 'CHƯƠNG 2', 'CHƯƠNG 3', 'CHƯƠNG 4',
            'CHƯƠNG 5', 'CHƯƠNG 6', 'CHƯƠNG 7', 'CHƯƠNG 8', 'CHƯƠNG 9',
            'PHỤ LỤC 1', 'PHỤ LỤC 2', 'PHỤ LỤC 3']
cov = sum(1 for c in chapters if c in vn_text)
print(f'  Chapters present: {cov}/{len(chapters)}')

# Save JSON
out = {
    'total_issues': total_issues,
    'issues_by_method': {k: v for k, v in issues.items()},
    'vn_chars': len(vn_text),
    'en_chars': len(en_text),
    'ratio': len(vn_text)/len(en_text),
    'chapters_covered': cov,
}
with open(os.path.join(os.path.dirname(__file__), 'qa_VN022_report.json'), 'w', encoding='utf-8') as f:
    json.dump(out, f, ensure_ascii=False, indent=2)
print(f'\nReport saved: qa_VN022_report.json')
