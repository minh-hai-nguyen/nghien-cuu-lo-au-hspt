# -*- coding: utf-8 -*-
"""
RAG v6 — MULTI-SOURCE: report v3 + all 60 summaries + key translations.

Cai tien:
- Index TAT CA nguon (report v3 + 60 tom tat + 60 ban dich) voi metadata source
- Metadata: canonical_id, doc_type (report/summary/translation), section
- Multilingual embedding (fallback ve paraphrase-multilingual-MiniLM-L12-v2)
- Auto-extract canonical ID tu filename
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import chromadb
from sentence_transformers import SentenceTransformer

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DB_PATH = os.path.join(ROOT, 'rag_bao_cao_can_thiep')

# ============================================================
# 1. Collect all source documents
# ============================================================
sources = []

# Report v3
report_path = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 11042026 v3.docx')
sources.append(('REPORT_v3', report_path, 'report'))

# All 60 summaries
tt_dir = os.path.join(ROOT, 'Tom-tat-tung-bai')
for f in sorted(os.listdir(tt_dir)):
    if f.endswith('.docx') and (f.startswith('VN') or f.startswith('QT')):
        m = re.match(r'(VN\d{3}|QT\d{3})', f)
        if m:
            cid = m.group(1)
            sources.append((cid, os.path.join(tt_dir, f), 'summary'))

# All 60 translations
trans_dir = os.path.join(ROOT, '03_Ban-dich')
for f in sorted(os.listdir(trans_dir)):
    if f.endswith('.docx') and (f.startswith('VN') or f.startswith('QT')):
        m = re.match(r'(VN\d{3}|QT\d{3})', f)
        if m:
            cid = m.group(1)
            sources.append((cid, os.path.join(trans_dir, f), 'translation'))

print(f'Total sources: {len(sources)}')
print(f'  - Reports: {sum(1 for _, _, t in sources if t == "report")}')
print(f'  - Summaries: {sum(1 for _, _, t in sources if t == "summary")}')
print(f'  - Translations: {sum(1 for _, _, t in sources if t == "translation")}')

# ============================================================
# 2. Chunking
# ============================================================
def chunk_doc(path, canonical_id, doc_type):
    """Return list of chunks with context"""
    chunks = []
    try:
        doc = Document(path)
    except Exception as e:
        print(f'  ERR reading {path}: {e}')
        return chunks

    current_section = ''
    paras_buffer = []

    def flush_buffer():
        nonlocal paras_buffer
        if paras_buffer:
            text = '\n'.join(paras_buffer)
            if len(text) > 30:  # Skip very short
                chunks.append({
                    'text': text,
                    'section': current_section,
                })
            paras_buffer = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = para.style.name
        if style.startswith('Heading'):
            flush_buffer()
            current_section = t
            # Also add heading as its own chunk
            chunks.append({'text': t, 'section': current_section})
        else:
            paras_buffer.append(t)

    flush_buffer()

    # Tables as structured chunks
    for ti, tb in enumerate(doc.tables):
        if len(tb.rows) < 2:
            continue
        headers = [c.text.strip() for c in tb.rows[0].cells]
        parts = ['BẢNG: ' + ' | '.join(headers)]
        for row in tb.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            pairs = [f'{h}={v}' for h, v in zip(headers, cells) if v]
            if pairs:
                parts.append(' | '.join(pairs))
        chunks.append({'text': '\n'.join(parts), 'section': f'Table {ti+1}'})

    return chunks

all_chunks = []
for cid, path, dtype in sources:
    chunks = chunk_doc(path, cid, dtype)
    for i, c in enumerate(chunks):
        all_chunks.append({
            'id': f'{cid}_{dtype}_{i}',
            'text': c['text'],
            'canonical_id': cid,
            'doc_type': dtype,
            'section': c['section'],
            'source_file': os.path.basename(path),
        })
    if chunks:
        pass  # suppress verbose

print(f'\nTotal chunks: {len(all_chunks)}')

# ============================================================
# 3. Build embeddings + index
# ============================================================
print('\nLoading model paraphrase-multilingual-MiniLM-L12-v2...')
model = SentenceTransformer('paraphrase-multilingual-MiniLM-L12-v2')

client = chromadb.PersistentClient(path=DB_PATH)
try:
    client.delete_collection('rag_v6_multisource')
except Exception:
    pass

col = client.create_collection(
    'rag_v6_multisource',
    metadata={'hnsw:space': 'cosine',
              'description': 'RAG v6 multi-source: report v3 + 60 summaries + 60 translations'}
)

print(f'Encoding {len(all_chunks)} chunks...')
texts = [c['text'] for c in all_chunks]
ids = [c['id'] for c in all_chunks]
metas = [{'canonical_id': c['canonical_id'],
          'doc_type': c['doc_type'],
          'section': c['section'][:100],
          'source_file': c['source_file']} for c in all_chunks]

# Batch encode (avoid memory issues)
batch_size = 64
embs = []
for i in range(0, len(texts), batch_size):
    batch = texts[i:i+batch_size]
    e = model.encode(batch, show_progress_bar=False).tolist()
    embs.extend(e)
    if i % 256 == 0:
        print(f'  encoded {i+len(batch)}/{len(texts)}')

col.add(documents=texts, embeddings=embs, ids=ids, metadatas=metas)
print(f'\nAdded {col.count()} chunks to collection "rag_v6_multisource"')

# ============================================================
# 4. Smoke tests
# ============================================================
print('\n=== SMOKE TEST (10 queries) ===')

tests = [
    ('Happy House Tran 2023 hiệu quả trầm cảm Cohen d',
     ['Happy House', '0,11', '1.084', 'Tran']),
    ('Dong PLOS 2025 kênh giao tiếp gia đình OR 0,22',
     ['0,22', 'Dong', 'giao tiếp']),
    ('Walder DMHI SAD specific Hedges g 0,878',
     ['0,878', 'Walder', 'SAD']),
    ('Walkup CAMS NEJM đáp ứng CBT + SSRI',
     ['80,7', 'Walkup', 'CAMS']),
    ('Zhu 2025 Suzhou giấc ngủ AOR 13,71',
     ['13,71', 'Zhu', 'giấc ngủ']),
    ('VN030 Happy House RAP-A thích ứng văn hoá',
     ['RAP-A', 'thích ứng', 'Happy House']),
    ('QT040 Walder 21 RCT JMIR',
     ['Walder', '21 RCT', 'JMIR']),
    ('Wen 2020 OR 11,58 rural China',
     ['11,58', 'Wen', 'rural']),
    ('Xian NMA iCBT SUCRA 71,2',
     ['71,2', 'Xian', 'iCBT']),
    ('Tran Thao Vi 2024 Huế longitudinal ESSA',
     ['Thảo Vi', 'Huế', 'ESSA']),
]

passed = 0
for q, expected in tests:
    q_emb = model.encode([q]).tolist()
    res = col.query(query_embeddings=q_emb, n_results=5,
                    include=['documents', 'distances', 'metadatas'])
    top_text = ' '.join(res['documents'][0]).lower()
    found = [kw for kw in expected if kw.lower() in top_text]
    if len(found) >= 1:
        passed += 1
        status = 'PASS'
    else:
        status = 'FAIL'
    rel = (1 - res['distances'][0][0]) * 100
    top_cid = res['metadatas'][0][0].get('canonical_id', '')
    top_type = res['metadatas'][0][0].get('doc_type', '')
    print(f'[{status}] [{rel:.0f}%] [{top_cid}/{top_type}] "{q[:50]}" found={found}')

print(f'\nSmoke test: {passed}/{len(tests)} PASS ({passed*100//len(tests)}%)')

# ============================================================
# 5. Save manifest
# ============================================================
manifest = {
    'collection': 'rag_v6_multisource',
    'n_chunks': len(all_chunks),
    'n_sources': len(sources),
    'embedding_model': 'paraphrase-multilingual-MiniLM-L12-v2',
    'sources_by_type': {
        'report': sum(1 for _, _, t in sources if t == 'report'),
        'summary': sum(1 for _, _, t in sources if t == 'summary'),
        'translation': sum(1 for _, _, t in sources if t == 'translation'),
    },
    'smoke_test_pass_rate': f'{passed}/{len(tests)}',
}
with open(os.path.join(os.path.dirname(__file__), 'rag_v6_manifest.json'), 'w', encoding='utf-8') as f:
    json.dump(manifest, f, ensure_ascii=False, indent=2)
print(f'\nManifest saved.')
print(f'DB: {DB_PATH}')
