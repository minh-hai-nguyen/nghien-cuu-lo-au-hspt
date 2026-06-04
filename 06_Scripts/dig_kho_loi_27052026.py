# -*- coding: utf-8 -*-
"""Dig into suspicious findings from audit_kho."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

CHECKS = [
    # (filepath, list of suspicious strings to find context for)
    ('03_Ban-dich/VN016_BaoQuyen_2025_YHCD.docx', ['33,2%']),
    ('03_Ban-dich/VN018_AnGiang_2025_YHVN.docx', ['Lê Minh T.', 'Nguyễn Đăng', 'Long Bình']),
    ('Tom-tat-tung-bai/VN018_AnGiang_2025_YHVN.docx', ['Lê Minh T.', 'Nguyễn Đăng', '366 học sinh']),
    ('Tom-tat-tung-bai/QT005_Alharbi_et_al_2019_SaudiArabia.docx', ['25,8%', '63,4%']),
]

for relpath, kws in CHECKS:
    full = os.path.join(ROOT, relpath)
    print(f"\n=== {relpath} ===")
    if not os.path.exists(full):
        print(f"  KHONG TON TAI")
        continue
    d = Document(full)
    for kw in kws:
        print(f"\n  *** '{kw}' ***")
        # body
        for i, p in enumerate(d.paragraphs):
            if kw in p.text:
                print(f"    para {i}: {p.text[:300]}")
        # table
        for ti, tb in enumerate(d.tables):
            for ri, row in enumerate(tb.rows):
                for ci, cell in enumerate(row.cells):
                    if kw in cell.text:
                        print(f"    Table {ti+1} r{ri} c{ci}: {cell.text[:300]}")
