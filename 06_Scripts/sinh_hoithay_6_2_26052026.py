# -*- coding: utf-8 -*-
"""Sinh file Hoi thay ve cach sua Nhiem vu NC 6.2 - 26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '07_HoiThay_NhiemVu_6_2_v1_26052026.docx')

def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)

def P(doc, text, indent=True, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    r.italic = italic; r.bold = bold
    if color is not None:
        r.font.color.rgb = color
    return p

def P_quote(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    r = p.add_run('"' + text + '"')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.italic = True

def P_option(doc, label, old_text, new_text, note):
    """Đoạn trình bày 1 option."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    # Label đậm
    r = p.add_run(f'{label}: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    # Text mô tả
    r = p.add_run(note)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)

    # Đoạn 6.2 cũ
    p_old = doc.add_paragraph()
    p_old.paragraph_format.left_indent = Cm(1.5)
    p_old.paragraph_format.right_indent = Cm(0.5)
    r = p_old.add_run('• 6.2 hiện tại: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True
    r = p_old.add_run(old_text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True

    # Đoạn 6.2 sửa (chữ đỏ đậm cho phần thay đổi)
    p_new = doc.add_paragraph()
    p_new.paragraph_format.left_indent = Cm(1.5)
    p_new.paragraph_format.right_indent = Cm(0.5)
    r = p_new.add_run('• 6.2 đề xuất: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True; r.bold = True
    # New text with red bold for changed parts
    for segment, is_red in new_text:
        r = p_new.add_run(segment)
        r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = True
        if is_red:
            r.font.color.rgb = RGBColor(192, 0, 0)
            r.bold = True

# ============================================================
# BUILD
# ============================================================
doc = doc_init()

# Title
H(doc, 'HỎI THẦY VỀ CÁCH SỬA NHIỆM VỤ NGHIÊN CỨU 6.2', 1)
P(doc, '(NCS Công Thị Hằng — 26/05/2026)', indent=False, italic=True)
doc.add_paragraph()

# Greeting
P(doc, 'Kính gửi thầy,')
P(doc, 'Em xin cảm ơn thầy đã có thêm góp ý về cách sửa Nhiệm vụ nghiên cứu 6.2. Em hiểu ý thầy là phần "biện pháp đối phó với rối loạn lo âu" đứng riêng trong nhiệm vụ 6.2 thì hơi hẹp, không phản ánh được toàn bộ phạm vi của luận án. Em xin đề xuất ba phương án sửa để thầy chọn, cùng với ý kiến đề xuất của em ở cuối.')
doc.add_paragraph()

# Tin nhắn của thầy
P(doc, 'Trích tin nhắn thầy:', indent=False, italic=True)
P_quote(doc, 'Thầy chỉ muốn xem lại cách sửa Nhiệm vụ NC 6.2. Hay là thay bằng từ lý thuyết và thực trạng, đề xuất khung tập huấn. Chỉ đưa biện pháp đối phó nhỏ quá.')
doc.add_paragraph()

# Phương án A
H(doc, 'Phương án A — Mở rộng 6.2, giữ 6.1 và 6.3 riêng (em đề xuất ưu tiên)', 2)
P(doc, 'Giữ ba nhiệm vụ riêng biệt 6.1, 6.2, 6.3 như cấu trúc hiện tại; chỉ sửa nội dung 6.2 để (i) loại bỏ "biện pháp đối phó" như một mục riêng và đưa nó vào trong yếu tố bảo vệ; (ii) bổ sung scope "lý thuyết" và "đề xuất khung tập huấn" làm yếu tố làm rõ phạm vi nghiên cứu.')

P_option(
    doc,
    'Cụ thể',
    'Nghiên cứu thực trạng, biểu hiện mức độ rối loạn lo âu ở học sinh trung học cơ sở và các yếu tố ảnh hưởng (yếu tố nguy cơ, yếu tố bảo vệ), biện pháp đối phó với rối loạn lo âu; Nghiên cứu trường hợp điển hình.',
    [
        ('Nghiên cứu ', False),
        ('cơ sở lý thuyết và ', True),
        ('thực trạng rối loạn lo âu ở học sinh trung học cơ sở và các yếu tố ảnh hưởng (yếu tố nguy cơ, yếu tố bảo vệ ', False),
        ('trong đó có biện pháp đối phó', True),
        ('); ', False),
        ('đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu', True),
        ('; Nghiên cứu trường hợp điển hình.', False),
    ],
    ''
)
P(doc, 'Ưu điểm: Giữ được cấu trúc ba nhiệm vụ rõ ràng; chỉ chỉnh nhẹ 6.2 để vừa rộng vừa nhất quán với khung "ứng phó là yếu tố bảo vệ" như thầy gợi ý. Có một chút trùng lặp nhẹ với 6.1 (lý luận) và 6.3 (đề xuất khung), song trùng lặp này phổ biến trong cấu trúc nhiệm vụ luận án và không gây mâu thuẫn.')
doc.add_paragraph()

# Phương án B
H(doc, 'Phương án B — Gộp 6.2 + 6.3 thành một nhiệm vụ broader, bỏ 6.3', 2)
P(doc, 'Cấu trúc nhiệm vụ rút gọn còn hai mục: 6.1 (cơ sở lý luận) và 6.2 (thực trạng + đề xuất khung). Nội dung 6.3 cũ được tích hợp vào 6.2.')

P_option(
    doc,
    'Cụ thể',
    '6.2: Nghiên cứu thực trạng... biện pháp đối phó; Nghiên cứu trường hợp điển hình. — 6.3: Đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu học đường cho học sinh THCS.',
    [
        ('Nghiên cứu lý thuyết và thực trạng rối loạn lo âu ở học sinh trung học cơ sở, các yếu tố ảnh hưởng (yếu tố nguy cơ, yếu tố bảo vệ ', False),
        ('bao gồm biện pháp đối phó', True),
        ('). ', False),
        ('Đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh THCS', True),
        ('. Nghiên cứu trường hợp điển hình.', False),
    ],
    ''
)
P(doc, 'Ưu điểm: Cấu trúc gọn hai nhiệm vụ; không trùng lặp.')
P(doc, 'Nhược điểm: Phải sửa lại danh mục nhiệm vụ, xóa 6.3 — thay đổi sâu hơn.')
doc.add_paragraph()

# Phương án C
H(doc, 'Phương án C — Theo nghĩa đen lời thầy (chỉ thay cụm từ)', 2)
P(doc, 'Trong 6.2 chỉ thay phrase "biện pháp đối phó với rối loạn lo âu" bằng phrase "lý thuyết và thực trạng, đề xuất khung tập huấn", giữ cấu trúc câu nguyên bản.')

P_option(
    doc,
    'Cụ thể',
    'Nghiên cứu thực trạng, biểu hiện mức độ rối loạn lo âu ở học sinh trung học cơ sở và các yếu tố ảnh hưởng (yếu tố nguy cơ, yếu tố bảo vệ), biện pháp đối phó với rối loạn lo âu; Nghiên cứu trường hợp điển hình.',
    [
        ('Nghiên cứu thực trạng, biểu hiện mức độ rối loạn lo âu ở học sinh trung học cơ sở và các yếu tố ảnh hưởng (yếu tố nguy cơ, yếu tố bảo vệ ', False),
        ('bao gồm biện pháp đối phó', True),
        ('), ', False),
        ('lý thuyết và thực trạng, đề xuất khung tập huấn', True),
        ('; Nghiên cứu trường hợp điển hình.', False),
    ],
    ''
)
P(doc, 'Đây là cách dịch sát chữ lời thầy nhất, song câu kết quả hơi cồng kềnh về ngữ pháp do gắn nhiều cụm trong một danh mục.')
doc.add_paragraph()

# Đề xuất
H(doc, 'Đề xuất của em', 2)
P(doc, 'Theo em, Phương án A là cách diễn giải tự nhiên nhất từ ý của thầy và ít gây xáo trộn nhất cho cấu trúc nhiệm vụ. Phương án này: (i) loại bỏ "biện pháp đối phó" như một mục riêng ngoài YTBV — đáp ứng critique của phản biện về tính nhất quán biến; (ii) thêm "cơ sở lý thuyết" và "đề xuất khung tập huấn" vào scope 6.2 — làm rộng phạm vi nhiệm vụ trung tâm như thầy gợi ý; (iii) giữ ba nhiệm vụ 6.1, 6.2, 6.3 riêng biệt — không phải xáo trộn cấu trúc.')

P(doc, 'Tuy nhiên, nếu thầy thấy Phương án B (gộp 6.2 và 6.3) hợp hơn với ý của thầy, em xin sửa theo Phương án B. Em xin chờ ý kiến của thầy để cập nhật bản luận án.')

doc.add_paragraph()
P(doc, 'Em xin cảm ơn thầy.', indent=False)
P(doc, 'Trân trọng,', indent=False)
P(doc, '', indent=False)
P(doc, 'Công Thị Hằng', indent=False, bold=True)

# ============================================================
# SAVE
# ============================================================
doc.save(OUT)
from docx import Document as D
d = D(OUT)
w = sum(len(p.text.split()) for p in d.paragraphs)
n_p = sum(1 for p in d.paragraphs if p.text.strip())
print(f"Da luu: {OUT}")
print(f"Words: {w}, Paragraphs: {n_p}, Size: {os.path.getsize(OUT)//1024}KB")
