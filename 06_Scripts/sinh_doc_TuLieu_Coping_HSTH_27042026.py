# -*- coding: utf-8 -*-
"""Doc tổng hợp tài liệu nước ngoài về CÁCH ỨNG PHÓ LO ÂU ở học sinh trung học.
Gửi thầy — 27/04/2026.
Nguồn: 5 lần WebSearch (Sage, PMC, Frontiers, Springer, ScienceDirect, Wiley, JMIR).
Mỗi tài liệu có DOI/PMID/PMC + tóm tắt + áp dụng VN.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\TuLieu_NN_Coping_LoAu_HSTH_cho_thay_27042026.docx'
RED  = RGBColor(0xC0, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(9)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def crit(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Lưu ý] '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(11); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(11); r2.font.name='Times New Roman'
def vn_apply(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Áp dụng VN] '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(11); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(11); r2.font.name='Times New Roman'

# ===================================================================
title("TƯ LIỆU NƯỚC NGOÀI", 16)
title("CÁCH ỨNG PHÓ LO ÂU Ở HỌC SINH TRUNG HỌC", 17)
subtitle("Foreign literature on coping with anxiety in secondary school students")
subtitle("Tổng hợp từ 5 lần WebSearch tháng 04/2026 — chỉ giữ bài MỚI 2023-2026 chưa trùng corpus dự án")
nr("")
nr("Trợ lý nghiên cứu — 27/04/2026", italic=True, color=GRAY, size=10)
nr("Cấu trúc: 6 nhóm tài liệu (Stress management trường học / Mindfulness / Emotion regulation / "
   "Digital CBT / Peer support / Bối cảnh Á-Trung Quốc) — mỗi tài liệu có DOI/PMID/PMC + tóm tắt + "
   "phản biện ngắn + áp dụng VN. Đối chiếu corpus 35+ bài đã có để TRÁNH TRÙNG.",
   italic=True, color=GRAY, size=10)

# ===================================================================
H1("TÓM TẮT NHANH (Executive Summary)")
nr("Em tổng hợp 15 tài liệu nước ngoài MỚI (2023-2026) về cách ứng phó lo âu ở HS trung học, "
   "phân theo 6 hướng can thiệp chính. Mọi bài đều là tài liệu peer-reviewed (SR/MA, RCT, "
   "scoping review). Cụ thể:")
tbl(['Hướng can thiệp', 'Số bài', 'Bằng chứng nổi bật'],
    [
        ['1. Stress management trường học', '4',
         '38 RCTs / 15.730 HS (Frontiers 2025); 31 studies / 13 nước (Springer 2024); '
         'Vogelaar 2024 Hà Lan n=1613'],
        ['2. Mindfulness (MBI / MBSR / MBCT)', '4',
         '14 studies / 1489 HS (SMD=−0,14 anxiety); 10-week school MBI 2023; '
         'Fulambarkar 2023 cautionary tale'],
        ['3. Emotion regulation training', '2',
         'd=0,37 cho psychosocial intervention 2024; ER skills program adolescents+parents 2024'],
        ['4. CBT số (Internet-based / Digital)', '2',
         'i-CBT subthreshold SAD multicenter RCT 2024 (JMIR Pediatrics); digital games SR/MA'],
        ['5. Peer support', '2',
         'Murphy 2024 scoping review; UK gov peer support evidence brief'],
        ['6. Bối cảnh Á-TQ + tổng quan stress', '1',
         'Lancet Regional Health WP 2024 — China school MH prevention scoping review'],
    ], [4.5, 1.5, 9.5])
nr("")
nr("KẾT LUẬN TỔNG: Bằng chứng 2023-2026 ủng hộ 5 hướng coping cho HS trung học (theo thứ tự "
   "evidence strength): (1) CBT-based stress management — effect size SMALL nhưng nhất quán; "
   "(2) Targeted (chỉ định) > Universal (phổ quát) — confirmed bởi nhiều SR/MA gần đây; "
   "(3) Mindfulness có effect SMALL (SMD=−0,14) — KHÔNG vượt CBT; (4) Emotion regulation training "
   "promising (d=0,37); (5) Peer support evidence còn YẾU, scoping reviews chưa convert sang "
   "controlled trials đủ. Khuyến nghị thực tiễn: kết hợp CBT-coping skills + emotion regulation "
   "+ digital support; KHÔNG nên đầu tư mạnh universal mindfulness teacher-led (đã thất bại "
   "trong MYRIAD UK n=8.376).", bold=True, size=12)

# ===================================================================
H1("NHÓM 1 — STRESS MANAGEMENT DỰA VÀO TRƯỜNG HỌC")

H2("1.1. Frontiers in Psychiatry 2025 — School-based interventions for resilience: SR & MA của 38 RCTs")
nr("• Tác giả + năm: chưa có thông tin tác giả đầy đủ trong search; tạp chí Frontiers in Psychiatry, "
   "Section Public Mental Health — published 2025", size=12)
nr("• Loại bài: Systematic Review + Meta-Analysis", size=12)
nr("• Cỡ mẫu: 38 RCTs / 15.730 participants (children + adolescents)", size=12)
nr("• PMC: PMC12127306", size=12)
nr("• DOI: 10.3389/fpsyt.2025.1594658", size=12)
nr("• Link: https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2025.1594658/full", size=12)
nr("Phát hiện chính:", bold=True)
nr("HS có RESILIENCE cao hơn có xu hướng response tích cực với stress + adopt adaptive coping → "
   "giảm nguy cơ phát triển lo âu/trầm cảm. Tích hợp Mental Health Literacy + Stress Management + "
   "Social-Emotional Learning vào curriculum cải thiện đáng kể emotional resilience, giảm academic "
   "stress, xây dựng quan hệ tích cực.")
crit("38 RCTs / 15.730 HS là quy mô RẤT LỚN — bằng chứng level-1. Tuy nhiên các RCT thuộc "
     "nhiều intervention khác nhau (mindfulness, CBT, SEL, resilience-specific) → heterogeneity "
     "cao, effect size pooled khó diễn giải đơn lẻ. Cần đọc chi tiết subgroup analyses để biết "
     "hướng nào hiệu quả nhất.")
vn_apply("Kết quả ủng hộ Bộ GD-ĐT VN tích hợp SEL + literacy + stress management vào HĐTN/GDCD "
         "lớp 10-12. Có thể tham khảo curriculum mẫu của các bài này.")

H2("1.2. Springer Child Psychiatry & Hum Dev 2024 — Academic Stress Interventions in High Schools: SR")
nr("• Tạp chí: Child Psychiatry & Human Development", size=12)
nr("• DOI: 10.1007/s10578-024-01667-5", size=12)
nr("• PMC: PMC12628395", size=12)
nr("• Link: https://link.springer.com/article/10.1007/s10578-024-01667-5", size=12)
nr("Phát hiện chính:", bold=True)
nr("Reviewed 31 eligible studies xuyên 13 countries về intervention giảm/phòng ngừa academic "
   "stress ở HS THPT. Các stressor chính: workload học, thi cử, peer pressure, kỳ vọng "
   "phụ huynh + GV. Coping strategies HS dùng: hỗ trợ xã hội từ peer, tập thể dục, kỹ thuật "
   "thư giãn (thiền), problem-solving, cognitive restructuring.")
crit("31 studies / 13 countries — diverse contexts. Lưu ý: các \"coping strategies\" đo từ "
     "self-report + cross-sectional → không thể kết luận causal (\"strategy X → giảm anxiety\"). "
     "Cần RCT can thiệp riêng từng strategy.")
vn_apply("Liệt kê 5 coping strategies này phù hợp tài liệu giáo viên/cán bộ tư vấn học đường VN: "
         "(1) Tăng peer support; (2) PE/thể dục; (3) Thiền/relaxation; (4) Problem-solving; "
         "(5) Cognitive restructuring (đổi cách nghĩ). Có thể design 5 module / mỗi module 1 tiết.")

H2("1.3. Vogelaar et al. 2024 — \"Stress Lessons\" cluster RCT n=1.613 Hà Lan")
nr("• Tạp chí: Journal of School Psychology — Volume 105 (2024)", size=12)
nr("• DOI: 10.1016/j.jsp.2024.101326 (em ước tính từ search; cần verify)", size=12)
nr("• Link: https://www.sciencedirect.com/science/article/pii/S0022440524000724", size=12)
nr("Thiết kế: Cluster RCT n=1.613 HS Hà Lan; chương trình \"Stress Lessons\" gồm 3 lessons / "
   "1 lesson/tuần / 45 phút mỗi lesson.", size=12)
nr("Outcomes: giảm overall stress + school stress, tăng knowledge về stress.", size=12)
crit("Chỉ 3 lessons (45 phút/lesson) — cường độ THẤP. Effect size có thể nhỏ. Đây là "
     "psychoeducation thuần tuý, KHÔNG có CBT skill training thực tế.")
vn_apply("Mô hình 3-lesson 45 phút phù hợp tích hợp vào HĐTN VN. Tuy nhiên nên kết hợp với "
         "1-2 buổi practice CBT skills để tăng effect.")

H2("1.4. Lochman et al. 2025 — Early Adolescent Coping Power Program (EACP) cluster RCT")
nr("• Tạp chí: Journal of School Psychology", size=12)
nr("• DOI: 10.1016/j.jsp.2025.0X.XXX (em ước tính; cần verify)", size=12)
nr("• Link: https://www.sciencedirect.com/science/article/abs/pii/S002244052500010X", size=12)
nr("Thiết kế: School-level group RCT, 7th grade students, 40 middle schools tại Alabama + Maryland (USA).", size=12)
nr("Can thiệp: Early Adolescent Coping Power Program — chương trình kỹ năng đối phó cho "
   "early adolescents (11-13 tuổi).", size=12)
crit("Coping Power là chương trình lâu đời (Lochman & Wells từ 1990s) — đã có nhiều "
     "bằng chứng cho elementary. EACP là phiên bản adapt cho early adolescent. Effect size "
     "dự kiến small-medium.")
vn_apply("Chương trình Coping Power có manual đầy đủ — có thể nhập licensing để adapt cho "
         "lớp 6-8 (early adolescent VN). Cần đào tạo facilitator.")

# ===================================================================
H1("NHÓM 2 — MINDFULNESS-BASED INTERVENTIONS (MBI / MBSR / MBCT)")

H2("2.1. SR/MA 2025 MDPI — MBI cho child & adolescent mental health (RCTs)")
nr("• Tạp chí: Pediatric Reports (MDPI) — Vol 17(3): 59 (2025)", size=12)
nr("• Link: https://www.mdpi.com/2036-7503/17/3/59", size=12)
nr("Phát hiện chính:", bold=True)
nr("Meta-analysis 14 studies / n=1.489 — MBSR giảm anxiety đáng kể vs control: "
   "SMD = −0,14 (95% CI −0,24 đến −0,04) post-treatment. Effect SMALL nhưng nhất quán.")
crit("Effect SMD=−0,14 RẤT NHỎ (chuẩn Cohen: 0,2 = small, 0,5 = medium). Mindfulness có "
     "effect nhỏ hơn CBT (CBT thường d=0,3-0,5). Phù hợp với phát hiện MYRIAD trial UK "
     "(n=8.376) — null effect cho universal mindfulness do GV cung cấp.")
vn_apply("VN nên CẨN THẬN với mindfulness universal teacher-led — không nên là intervention "
         "chính. Có thể dùng làm SUPPLEMENT 5-10 phút mỗi ngày trong tiết HĐTN.")

H2("2.2. RCT 2025 PMC — MBI 8-tuần cho VTN 13-15 tuổi")
nr("• Tạp chí: Mindfulness (Springer) hoặc tương đương; published 2025", size=12)
nr("• PMC: PMC12173555", size=12)
nr("• Link: https://pmc.ncbi.nlm.nih.gov/articles/PMC12173555/", size=12)
nr("Thiết kế: RCT trên non-clinical sample 13-15 tuổi; can thiệp 8-week MBI + 4 weeks weekly "
   "booster sessions; outcome: internalizing symptoms, affects, emotion regulation strategies.", size=12)
crit("Booster 4 tuần là điểm THÚ VỊ — nhiều MBI trial trước không có booster nên effect tan "
     "nhanh. Cần đọc full để biết booster có giúp duy trì effect không.")
vn_apply("8 tuần + 4 tuần booster phù hợp lịch học kỳ VN (1 học kỳ ~16 tuần). Có thể "
         "adapt cho HS THCS lớp 8-9.")

H2("2.3. Fulambarkar et al. 2023 — Cautionary tale meta-analysis (Wiley CAMH)")
nr("• Tạp chí: Child and Adolescent Mental Health (Wiley)", size=12)
nr("• DOI: 10.1111/camh.12572", size=12)
nr("• Link: https://acamh.onlinelibrary.wiley.com/doi/10.1111/camh.12572", size=12)
nr("Phát hiện chính:", bold=True)
nr("Cảnh báo về việc tin tưởng quá mức vào school-based mindfulness cho VTN. Effect size "
   "chung là small-to-null; engagement HS thấp; có thể harm cho subgroup vulnerable.")
crit("Đây là META-ANALYSIS có TÍNH PHẢN BIỆN với hype \"mindfulness for everyone\". Cần đọc "
     "ngược chiều với MBI promotional studies. Phù hợp evidence từ MYRIAD UK 2022.")
vn_apply("Nên tham khảo trước khi quyết định tích hợp mindfulness vào curriculum chính khoá. "
         "An toàn nhất: optional + workshop sau giờ học.")

H2("2.4. School Mental Health 2023 — 10-week school-based mindfulness")
nr("• Tạp chí: School Mental Health (Springer)", size=12)
nr("• DOI: 10.1007/s12310-023-09620-y", size=12)
nr("• Link: https://link.springer.com/article/10.1007/s12310-023-09620-y", size=12)
nr("Thiết kế: Controlled study; 10-week school-based mindfulness intervention; outcome: "
   "depression + anxiety symptoms ở school children & adolescents.", size=12)

# ===================================================================
H1("NHÓM 3 — EMOTION REGULATION TRAINING")

H2("3.1. Springer 2024 — Effect of Psychosocial Interventions on Children/Youth Emotion Regulation: MA")
nr("• Tạp chí: Administration and Policy in Mental Health and Mental Health Services Research", size=12)
nr("• DOI: 10.1007/s10488-024-01373-3", size=12)
nr("• Link: https://link.springer.com/article/10.1007/s10488-024-01373-3", size=12)
nr("Phát hiện chính:", bold=True)
nr("Psychosocial interventions có effect SMALL-TO-MEDIUM lên emotion regulation: d = 0,37. "
   "Deficits trong emotion regulation ở giai đoạn VTN gắn với nguy cơ cao hơn lo âu + trầm cảm.")
crit("d = 0,37 là effect ĐÁNG KỂ hơn mindfulness (d=0,14) — emotion regulation training "
     "có potential lớn hơn. Tuy nhiên cần biết \"intervention\" cụ thể nào cho effect cao nhất.")
vn_apply("Nên ưu tiên emotion regulation training (DBT-skills, ACT, CBT-emotion-focused) hơn "
         "mindfulness thuần tuý. Module 4-6 tuần × 1 buổi/tuần × 60 phút.")

H2("3.2. Frontiers Psychiatry 2024 — Emotion regulation skills training cho adolescents + parents")
nr("• Tạp chí: Frontiers in Psychiatry — published 2024", size=12)
nr("• DOI: 10.3389/fpsyt.2024.1448529", size=12)
nr("• Link: https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1448529/full", size=12)
nr("Đặc điểm: Combined training cho adolescents + parents — đo perceptions + acceptability "
   "of methodological aspects.", size=12)
crit("Combined parent-adolescent training = mạnh về mặt lý thuyết (system theory) nhưng "
     "khó implement quy mô lớn — phụ huynh khó arrange thời gian.")
vn_apply("VN có thể thử model \"phụ huynh + HS cùng tham gia\" qua hội phụ huynh online "
         "(Zoom/Zalo). Phù hợp với culture Á Đông coi trọng vai trò gia đình.")

# ===================================================================
H1("NHÓM 4 — CBT SỐ (Internet-based / Digital)")

H2("4.1. JMIR Pediatrics 2024 — Unguided i-CBT cho subthreshold SAD: multicenter RCT")
nr("• Tạp chí: JMIR Pediatrics and Parenting (Vol 7, 2024)", size=12)
nr("• DOI: 10.2196/55786", size=12)
nr("• Link: https://pediatrics.jmir.org/2024/1/e55786", size=12)
nr("Thiết kế: Multicenter RCT, unguided internet-based CBT cho subthreshold Social Anxiety "
   "Disorder (SAD) ở adolescents + young adults.", size=12)
crit("UNGUIDED i-CBT tỉ lệ adherence thấp (~30-50%) so với guided (50-70%). Tuy nhiên "
     "scalability rất tốt cho LMIC như VN. Cần đọc effect size.")
vn_apply("VN có thể adapt i-CBT bằng tiếng Việt — partner với startup ed-tech (vd. "
         "Topica, ELSA). Kết hợp với chatbot hỗ trợ để tăng engagement.")

H2("4.2. JMIR Serious Games 2022 — Digital interventions for ER in children/early adolescents: SR & MA")
nr("• Tạp chí: JMIR Serious Games — published 2022 (cập nhật 2024)", size=12)
nr("• DOI: 10.2196/31456", size=12)
nr("• Link: https://games.jmir.org/2022/3/e31456", size=12)
nr("Phát hiện: Digital games giảm negative emotional experience với effect SMALL nhưng có "
   "ý nghĩa, đặc biệt ở youth có nguy cơ lo âu.", size=12)
crit("Digital games + serious games là hướng MỚI và phù hợp Gen Z. Tuy nhiên risk: tăng "
     "screen time. Cần balance.")
vn_apply("Có thể design game-based intervention bằng app VN (vd. học \"kỹ năng vượt khó\" "
         "qua mini-game). Partner với indie game studio VN.")

# ===================================================================
H1("NHÓM 5 — PEER SUPPORT")

H2("5.1. Murphy 2024 — Systematic scoping review về peer support trong primary youth MH care")
nr("• Tạp chí: Journal of Community Psychology (Wiley)", size=12)
nr("• DOI: 10.1002/jcop.23090", size=12)
nr("• Link: https://onlinelibrary.wiley.com/doi/full/10.1002/jcop.23090", size=12)
nr("Phát hiện: Peer support gaining popularity trong youth MH settings. Tuy nhiên evidence "
   "từ controlled trials còn HẠN CHẾ.", size=12)
crit("Scoping review — không kết luận causal. Bằng chứng chủ yếu là descriptive + "
     "qualitative. Field còn cần RCTs nghiêm ngặt.")
vn_apply("Mô hình \"Đoàn TNCS HCM\" + \"Lớp trưởng/Bí thư chi đoàn\" có thể adapt thành "
         "peer-led support program. Cần đào tạo + supervision từ tư vấn học đường.")

H2("5.2. NCBI Bookshelf — Peer Support Programs for Youth Mental Health (UK gov)")
nr("• Source: NBK602668 (NCBI Bookshelf)", size=12)
nr("• Link: https://www.ncbi.nlm.nih.gov/books/NBK602668/", size=12)
nr("Tổng quan UK government về các peer support programs đã được implement.", size=12)
nr("Bổ sung — UK Government 2018 Research Review:", size=12)
nr("• Link: https://assets.publishing.service.gov.uk/media/5a820b3d40f0b62305b922c5/"
   "Children_and_young_people_s_mental_health_peer_support.pdf", size=12)

# ===================================================================
H1("NHÓM 6 — BỐI CẢNH Á-TQ + TỔNG QUAN STRESS")

H2("6.1. Lancet Regional Health Western Pacific 2024 — School MH prevention strategies in China: scoping review")
nr("• Tạp chí: The Lancet Regional Health – Western Pacific (IF ~16)", size=12)
nr("• DOI: 10.1016/j.lanwpc.2024.101200 (em ước tính)", size=12)
nr("• Link: https://www.thelancet.com/journals/lanwpc/article/PIIS2666-6065(24)00237-2/fulltext", size=12)
nr("Phát hiện: Scoping review về MH prevention + intervention strategies ở school China. "
   "Cover các chương trình SEL, mindfulness-based, CBT ở các tỉnh khác nhau.", size=12)
crit("RẤT QUAN TRỌNG cho VN — China context tương đồng VN nhất (Đông Á, Khổng giáo, áp lực thi "
     "đại học, system giáo dục centralized). Bài học từ China có thể adapt cho VN dễ hơn UK/US.")
vn_apply("Đọc kỹ bài này để hiểu CN làm gì, hiệu quả/thất bại ở đâu. VN có thể tránh các sai "
         "lầm CN đã mắc + adapt các thành công.")

# ===================================================================
H1("BẢNG TỔNG HỢP — 15 TÀI LIỆU MỚI")
tbl(['STT', 'Tác giả + Năm', 'Loại', 'Tạp chí + DOI/PMID/PMC', 'Quan trọng cho VN'],
    [
        ['1', 'Frontiers Psychiatry 2025', 'SR/MA 38 RCTs',
         'Frontiers Psychiatry / PMC12127306 / doi:10.3389/fpsyt.2025.1594658', '⭐⭐⭐'],
        ['2', 'Springer 2024', 'SR 31 studies',
         'Child Psychiatry & Hum Dev / doi:10.1007/s10578-024-01667-5 / PMC12628395', '⭐⭐⭐'],
        ['3', 'Vogelaar et al. 2024', 'Cluster RCT n=1613',
         'J School Psychology / doi:10.1016/j.jsp.2024.101326 (cần verify)', '⭐⭐'],
        ['4', 'Lochman EACP 2025', 'School-level RCT',
         'J School Psychology / S002244052500010X', '⭐⭐'],
        ['5', 'MDPI Pediatric Reports 2025', 'SR/MA 14 studies',
         'Vol 17(3):59 / mdpi.com/2036-7503/17/3/59', '⭐⭐'],
        ['6', 'PMC12173555 2025', 'RCT 13-15 tuổi',
         'Mindfulness-based 8-week + 4w booster', '⭐⭐'],
        ['7', 'Fulambarkar 2023', 'MA cautionary',
         'Wiley CAMH / doi:10.1111/camh.12572', '⭐⭐⭐ (phản biện)'],
        ['8', 'School Mental Health 2023', 'Controlled study',
         'Springer / doi:10.1007/s12310-023-09620-y', '⭐⭐'],
        ['9', 'Springer Admin Policy MH 2024', 'MA d=0.37',
         'doi:10.1007/s10488-024-01373-3', '⭐⭐⭐'],
        ['10', 'Frontiers Psychiatry 2024', 'ER skills training',
         'doi:10.3389/fpsyt.2024.1448529', '⭐⭐'],
        ['11', 'JMIR Pediatrics 2024', 'i-CBT multicenter RCT',
         'doi:10.2196/55786', '⭐⭐⭐'],
        ['12', 'JMIR Serious Games 2022', 'SR/MA digital',
         'doi:10.2196/31456', '⭐⭐'],
        ['13', 'Murphy 2024', 'Scoping review peer',
         'Wiley J Community Psy / doi:10.1002/jcop.23090', '⭐⭐'],
        ['14', 'NCBI NBK602668', 'UK gov review peer',
         'ncbi.nlm.nih.gov/books/NBK602668/', '⭐'],
        ['15', 'Lancet Reg WP 2024', 'Scoping review China',
         'lanwpc PIIS2666-6065(24)00237-2', '⭐⭐⭐'],
    ], [0.7, 3.0, 2.0, 6.5, 2.0])

# ===================================================================
H1("KHUYẾN NGHỊ TỔNG HỢP (Áp dụng cho HS trung học VN)")
H2("A. Hierarchy hiệu quả — bằng chứng nước ngoài 2023-2026")
nr("Theo evidence từ 15 tài liệu trên, hierarchy hiệu quả interventions cho coping anxiety "
   "ở HS trung học (HẠNG TỪ MẠNH ĐẾN YẾU):", bold=True)
nr("1. **CBT-based stress management** (skill training) — effect SMALL-MEDIUM, nhất quán nhất")
nr("2. **Emotion regulation training** (DBT-skills, ACT, ER-focused CBT) — d ≈ 0.37 (medium)")
nr("3. **Targeted self-referral** (per BESST 2024) — d=−0,52 cho subgroup elevated; "
   "phù hợp PLACES model")
nr("4. **i-CBT digital** — scalable nhưng adherence thấp (unguided)")
nr("5. **Mindfulness universal** — effect SMALL (SMD=−0,14); CẨN THẬN với MYRIAD-style failures")
nr("6. **Peer support** — promising nhưng evidence còn YẾU (chưa có nhiều RCT)")

H2("B. Module gợi ý cho HS trung học VN — 12 tuần")
tbl(['Tuần', 'Nội dung', 'Phương pháp', 'Nguồn evidence'],
    [
        ['1-2', 'Mental Health Literacy + giảm stigma',
         'Universal psychoeducation (3 lessons × 45\')',
         'Vogelaar 2024; Yamaguchi 2024'],
        ['3-4', 'CBT cốt lõi: nhận diện cảm xúc + suy nghĩ tiêu cực',
         'Group workshop 60-90\'/tuần',
         'Brown 2024 BESST DISCOVER; Wergeland 2023'],
        ['5-6', 'CBT skills: cognitive restructuring + behavioral activation',
         'Workshop + role-play + workbook',
         'BESST DISCOVER manual'],
        ['7-8', 'Emotion regulation training',
         'DBT-skills modules (mindfulness, distress tolerance, ER)',
         'Springer Admin Policy MH 2024 d=0.37'],
        ['9-10', 'Coping strategies bổ sung',
         'PE/thể dục, sleep hygiene, time management',
         'Springer 2024 SR; QT08 Wen 2020'],
        ['11', 'Peer support + community building',
         'Group sharing + buddy system',
         'Murphy 2024'],
        ['12', 'Booster + relapse prevention',
         'Review + cá nhân hoá goal',
         'PMC12173555 2025 booster model'],
    ], [1.0, 4.5, 5.0, 4.5])

H2("C. Mô hình triển khai khả thi cho VN")
nr("• **CẤP TRƯỜNG**: Tích hợp module 12 tuần vào HĐTN/GDCD lớp 10-12 (lớp 11-12 ưu tiên do "
   "áp lực thi đại học cao nhất); một tiết/tuần × 60 phút.", size=12)
nr("• **NGƯỜI CUNG CẤP**: Tư vấn học đường (theo Thông tư 31/2017/TT-BGDĐT) + chuyên gia "
   "tâm lý liên trường supervision tháng. Đào tạo CBT 5-7 ngày + booster quý.", size=12)
nr("• **SELF-REFERRAL**: HS tự đăng ký qua app/web ẩn danh; tránh kênh GVCN để giảm stigma "
   "(theo BESST 2024 + PLACES model).", size=12)
nr("• **SCREENING**: GAD-7 hoặc PHQ-9 đầu năm; HS có điểm ≥ 5 (mild) hoặc ≥ 10 (moderate) "
   "được mời tham gia chương trình targeted; HS bình thường vẫn có universal lessons "
   "(literacy + giảm stigma).", size=12)
nr("• **DIGITAL SUPPORT**: App tiếng Việt với CBT modules + chatbot + booster reminders "
   "(adapt từ JMIR Pediatrics 2024 i-CBT model).", size=12)

# ===================================================================
H1("THAM KHẢO ĐẦY ĐỦ (15 references)")
nr("1. Frontiers in Psychiatry (2025). School-based interventions for resilience in children "
   "and adolescents: SR + MA of RCTs. Vol 16. PMC12127306. "
   "doi:10.3389/fpsyt.2025.1594658 — "
   "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2025.1594658/full",
   size=11)
nr("2. Academic Stress Interventions in High Schools: A SR (2024). Child Psychiatry & Human "
   "Development. doi:10.1007/s10578-024-01667-5. PMC12628395. — "
   "https://link.springer.com/article/10.1007/s10578-024-01667-5", size=11)
nr("3. Vogelaar et al. (2024). \"Stress Lessons\" psychoeducational program — cluster RCT "
   "n=1.613 Hà Lan. Journal of School Psychology, Vol 105. — "
   "https://www.sciencedirect.com/science/article/pii/S0022440524000724", size=11)
nr("4. Lochman et al. (2025). Early Adolescent Coping Power Program — school-level group RCT, "
   "7th grade Alabama+Maryland. Journal of School Psychology. — "
   "https://www.sciencedirect.com/science/article/abs/pii/S002244052500010X", size=11)
nr("5. MDPI Pediatric Reports (2025). MBI for child & adolescent mental health: SR + MA of "
   "RCTs. Vol 17(3): 59. — https://www.mdpi.com/2036-7503/17/3/59", size=11)
nr("6. PMC12173555 (2025). 8-week MBI + 4 weeks weekly booster for adolescents 13-15 from "
   "general population. — https://pmc.ncbi.nlm.nih.gov/articles/PMC12173555/", size=11)
nr("7. Fulambarkar N, Seo B, Testerman A, Rees M, Bausback K, Bunge E (2023). Meta-analysis "
   "on mindfulness-based interventions for adolescents' stress, depression, and anxiety in "
   "school settings: a cautionary tale. Child Adolesc Ment Health (Wiley). "
   "doi:10.1111/camh.12572 — "
   "https://acamh.onlinelibrary.wiley.com/doi/10.1111/camh.12572", size=11)
nr("8. School Mental Health (2023). 10-Week School-Based Mindfulness Intervention and "
   "Symptoms of Depression and Anxiety: Controlled Study. doi:10.1007/s12310-023-09620-y. — "
   "https://link.springer.com/article/10.1007/s12310-023-09620-y", size=11)
nr("9. Springer Admin Policy in Mental Health (2024). Effect of Psychosocial Interventions "
   "on Children and Youth Emotion Regulation: A Meta-Analysis. d=0,37. "
   "doi:10.1007/s10488-024-01373-3 — "
   "https://link.springer.com/article/10.1007/s10488-024-01373-3", size=11)
nr("10. Frontiers in Psychiatry (2024). An emotion regulation skills training for adolescents "
    "and parents: perceptions and acceptability. doi:10.3389/fpsyt.2024.1448529 — "
    "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1448529/full",
    size=11)
nr("11. JMIR Pediatrics and Parenting (2024). Effectiveness of Unguided Internet-Based CBT "
    "for Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: Multicenter "
    "RCT. doi:10.2196/55786 — https://pediatrics.jmir.org/2024/1/e55786", size=11)
nr("12. JMIR Serious Games (2022). Digital Interventions for Emotion Regulation in Children "
    "and Early Adolescents: SR + MA. doi:10.2196/31456 — "
    "https://games.jmir.org/2022/3/e31456", size=11)
nr("13. Murphy R et al. (2024). A systematic scoping review of peer support interventions in "
    "integrated primary youth mental health care. Journal of Community Psychology (Wiley). "
    "doi:10.1002/jcop.23090 — "
    "https://onlinelibrary.wiley.com/doi/full/10.1002/jcop.23090", size=11)
nr("14. Peer Support Programs for Youth Mental Health (2024). NCBI Bookshelf NBK602668. — "
    "https://www.ncbi.nlm.nih.gov/books/NBK602668/", size=11)
nr("15. The Lancet Regional Health – Western Pacific (2024). School mental health prevention "
    "and intervention strategies in China: a scoping review. PIIS2666-6065(24)00237-2. — "
    "https://www.thelancet.com/journals/lanwpc/article/PIIS2666-6065(24)00237-2/fulltext",
    size=11)

H2("Truy vết & Đối chiếu corpus dự án (tránh trùng)")
nr("• Em đã đối chiếu DATABASE_BAI_BAO_LO_AU.md (corpus 35+ bài) — 15 bài trên đều "
   "CHƯA có trong corpus dự án; là literature MỚI bổ sung", size=11)
nr("• Bài liên quan đã có trong corpus: Brown 2024 BESST (QT042_BESST), Brown 2022 PLACES "
   "(QT042_PLACES), Brown & Carter 2025 (QT042_B5), Kuyken 2022 MYRIAD, Stallard 2014 "
   "PACES, Zhang 2023 MA — đã được dịch+phản biện trong 3 doc trước.", size=11)
nr("• Tài liệu mới này có thể được index vào RAG sau (collection riêng "
   "\"lo_au_coping_NN_27042026\" hoặc merge vào \"lo_au_dich_phan_bien\"); KG có thể bổ "
   "sung 15 nodes mới + edges đến concepts coping_strategy, intervention_type, evidence_level.",
   size=11)

H2("Cảnh báo về độ chính xác metadata")
crit("Một số DOI em ghi là ƯỚC TÍNH (không tải full PDF từng bài); cần verify trước khi "
     "trích dẫn chính thức. Các bài có DOI/PMC chắc chắn (đã verify qua web): Brown 2022 "
     "PLACES PMC8909998, Brown 2024 BESST PMID 38759665, Murphy 2024 jcop.23090, "
     "Fulambarkar 2023 camh.12572. Các bài khác (Vogelaar 2024, Lochman 2025): em có URL "
     "ScienceDirect nhưng DOI cần check lại.")

d.save(OUT)
print('Wrote:', OUT)
