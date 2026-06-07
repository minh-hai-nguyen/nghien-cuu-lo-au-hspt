"""Extract facts/numbers from Q1 v7 and Q2 v1 for polish fact-check verification."""
import re
import json
from pathlib import Path
from docx import Document

ROOT = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au")
Q1V7 = ROOT / "bai-bao-Q1" / "Draft_Q1_SongNgu_v7_01062026.docx"
Q2V1 = ROOT / "bai-bao-Q1" / "Draft_Q2_SongNgu_v1_07062026.docx"


def read_all_text(path: Path) -> str:
    doc = Document(str(path))
    chunks = []
    for p in doc.paragraphs:
        if p.text.strip():
            chunks.append(p.text)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if p.text.strip():
                        chunks.append(p.text)
    return "\n".join(chunks)


def extract_facts(text: str) -> dict:
    facts = {}
    # Sample sizes N=...
    facts["N_values"] = sorted(set(re.findall(r"\bN\s*=\s*([\d.,]+)", text)))
    # Beta coefficients (β = ...)  also handle 'β=' without space and unicode
    beta = re.findall(r"[βΒbB]\s*=\s*(-?[\d.,]+)", text)
    facts["beta_values"] = sorted(set(beta))
    # p-values
    facts["p_values"] = sorted(set(re.findall(r"p\s*[<>=]\s*\.?\d+\.?\d*", text)))
    # alpha (Cronbach's) - α = 0.xx
    alpha = re.findall(r"[αΑ]\s*=\s*([\d.,]+)", text)
    facts["alpha_values"] = sorted(set(alpha))
    # CFA fit indices
    facts["CFI"] = sorted(set(re.findall(r"CFI\s*=\s*([\d.,]+)", text)))
    facts["TLI"] = sorted(set(re.findall(r"TLI\s*=\s*([\d.,]+)", text)))
    facts["RMSEA"] = sorted(set(re.findall(r"RMSEA\s*=\s*([\d.,]+)", text)))
    facts["SRMR"] = sorted(set(re.findall(r"SRMR\s*=\s*([\d.,]+)", text)))
    # Effect sizes (Cohen d / η2 / R2)
    facts["d_values"] = sorted(set(re.findall(r"\bd\s*=\s*(-?[\d.,]+)", text)))
    facts["R2_values"] = sorted(set(re.findall(r"R[²2]\s*=\s*([\d.,]+)", text)))
    # Confidence intervals
    facts["CI_values"] = sorted(set(re.findall(r"(?:95%\s*CI|CI\s*95%)\s*[=:\[]?\s*\[?(-?[\d.,]+\s*,\s*-?[\d.,]+)", text)))
    # Years (4 digit)
    facts["years"] = sorted(set(re.findall(r"\b(19\d{2}|20[0-2]\d)\b", text)))
    # Percentages / prevalence
    facts["pct_values"] = sorted(set(re.findall(r"\b(\d+(?:\.\d+)?)\s*%", text)))
    # χ² values
    facts["chi2"] = sorted(set(re.findall(r"χ[²2]\s*[=/]?\s*([\d.,]+)", text)))
    return facts


def extract_authors(text: str) -> list:
    # Capture cited author surnames pattern: "Surname (YEAR)" or "(Surname, YEAR)"
    patterns = [
        r"\b([A-Z][a-zA-ZàáạảãâầấậẩẫăằắặẳẵèéẹẻẽêềếệểễìíịỉĩòóọỏõôồốộổỗơờớợởỡùúụủũưừứựửữỳýỵỷỹđÀ-ỸÀ-ɏ\-]{2,})\s*\((19\d{2}|20[0-2]\d)\)",
        r"\(([A-Z][a-zA-ZÀ-ỸÀ-ɏ\-]{2,})\s*(?:et al\.?\s*)?,\s*(19\d{2}|20[0-2]\d)\)",
        r"\(([A-Z][a-zA-ZÀ-ỸÀ-ɏ\-]{2,})\s*&\s*[A-Z][a-zA-ZÀ-ỸÀ-ɏ\-]{2,}\s*,\s*(19\d{2}|20[0-2]\d)\)",
    ]
    pairs = set()
    for pat in patterns:
        for m in re.findall(pat, text):
            pairs.add(tuple(m))
    return sorted(pairs)


def check_specific_refs(text: str) -> dict:
    return {
        "Niwenahisemo_Hong_S": bool(re.search(r"Hong,\s*S", text)),
        "Niwenahisemo_Su_H_BAD": bool(re.search(r"Su,\s*H", text)),
        "Karasu_3094389": "3094389" in text,
        "Rose_12487497": "12487497" in text,
        "Stankov_DOI": "10.1016/j.lindif.2010.05.003" in text,
        "Small_Blanc_DOI": "10.3389/fpsyt.2020.589618" in text,
        "Anderson_T_L": bool(re.search(r"Anderson,?\s*T\.\s*L\.", text)),
        "DOI_count": len(re.findall(r"10\.\d{4,}/", text)),
        "PMID_count": len(re.findall(r"PMID:?\s*\d+", text, re.IGNORECASE)),
    }


def count_paragraphs(path: Path) -> dict:
    doc = Document(str(path))
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    return {
        "n_paragraphs": len(paras),
        "n_tables": len(doc.tables),
        "n_chars": sum(len(p) for p in paras),
    }


def main():
    out = {}
    for label, path in [("Q1v7", Q1V7), ("Q2v1", Q2V1)]:
        text = read_all_text(path)
        out[label] = {
            "stats": count_paragraphs(path),
            "facts": extract_facts(text),
            "authors": extract_authors(text)[:60],  # cap for readability
            "specific_refs": check_specific_refs(text),
            "text_len": len(text),
        }

    out_path = ROOT / "06_Scripts" / "_verify_q2polish_extract.json"
    with open(out_path, "w", encoding="utf-8") as f:
        json.dump(out, f, ensure_ascii=False, indent=2)
    print(f"Wrote {out_path}")

    # Diff facts
    diff_report = {}
    for key in out["Q1v7"]["facts"]:
        s1 = set(out["Q1v7"]["facts"][key])
        s2 = set(out["Q2v1"]["facts"][key])
        only1 = sorted(s1 - s2)
        only2 = sorted(s2 - s1)
        if only1 or only2:
            diff_report[key] = {"only_in_Q1v7": only1, "only_in_Q2v1": only2}
    diff_path = ROOT / "06_Scripts" / "_verify_q2polish_diff.json"
    with open(diff_path, "w", encoding="utf-8") as f:
        json.dump(diff_report, f, ensure_ascii=False, indent=2)
    print(f"Wrote {diff_path}")


if __name__ == "__main__":
    main()
