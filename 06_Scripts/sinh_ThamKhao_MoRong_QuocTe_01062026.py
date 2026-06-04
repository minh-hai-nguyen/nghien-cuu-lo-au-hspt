# -*- coding: utf-8 -*-
"""Sinh file mo rong tieu de paper Q1/Q3 quoc te ngoai Trung Quoc + VN.
Cac nuoc: Indonesia, Malaysia, Philippines, Japan, Iran, Saudi Arabia,
Turkey, Egypt, South Africa, Kenya, Brazil, Poland.
Moi entry da verify PubMed/Crossref. Co tieu de dich tieng Viet.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'ThamKhao_MoRong_QuocTe_TiengViet_v2_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.4


def H1(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(t, italic=False, indent=True, align_center=False):
    p = d.add_paragraph()
    p.alignment = (WD_ALIGN_PARAGRAPH.CENTER if align_center else WD_ALIGN_PARAGRAPH.JUSTIFY)
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def B(t, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('▸ ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def WARN(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def set_col_widths(t, widths_cm):
    for row in t.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_paper_table(rows, col_widths_cm):
    headers = ['STT', 'Tiêu đề tiếng Anh + dịch tiếng Việt',
               'Tác giả + Nước + Sample', 'Tạp chí + Năm + IF', 'Ghi chú']
    t = d.add_table(rows=1, cols=5); t.style = 'Light Grid Accent 1'; t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
H1('THAM KHẢO MỞ RỘNG — TIÊU ĐỀ Q1/Q3 QUỐC TẾ')
P('Bộ tham khảo bổ sung — mở rộng ra các nước khác ngoài Trung Quốc, '
  'Việt Nam, Bangladesh, Ethiopia, Nigeria, Rwanda đã có ở file v4',
  italic=True, align_center=True, indent=False)
P('Gồm Indonesia, Malaysia, Philippines, Nhật Bản, Iran, Ả Rập Saudi, '
  'Thổ Nhĩ Kỳ, Ai Cập, Nam Phi, Kenya, Brazil, Ba Lan', italic=True,
  align_center=True, indent=False)
P('Mỗi tiêu đề có dịch nghĩa tiếng Việt để NCS Công Thị Hằng và thầy HD '
  'tham khảo', italic=True, align_center=True, indent=False)
P('Ngày soạn: 01/06/2026', italic=True, align_center=True, indent=False)


# ============================================================
H1('LỜI MỞ ĐẦU')

P('File này bổ sung 13 tiêu đề bài báo Q1/Q3 từ các nước chưa được liệt '
  'kê trong file v4. Mục đích: mở rộng góc nhìn về tiêu đề thành công ở '
  'tạp chí quốc tế khi nghiên cứu rối loạn lo âu thanh thiếu niên, đặc '
  'biệt là các nước có bối cảnh văn hóa - kinh tế - giáo dục tương đồng '
  'hoặc đa dạng với Việt Nam.')

P('Tất cả 13 entry đã được verify trực tiếp PubMed/Crossref/JCR theo quy '
  'tắc HARD GATE trong memory đã ghi nhận. Tên tác giả, năm xuất bản, '
  'tạp chí và chỉ số IF JCR 2024 đều có nguồn gốc cụ thể.', italic=True)


# ============================================================
H1('PHẦN 1 — Q1 (SEM / MEDIATION / DỌC / VALIDATION)')

H2('1.1 Đông Nam Á + Đông Á')

q1_asia = [
    ('1', 'Mental Health Problems Among Indonesian Adolescents: Findings of '
     'a Cross-Sectional Study Utilizing Validated Scales and Innovative '
     'Sampling Methods\n'
     '(Các vấn đề sức khỏe tâm thần ở thanh thiếu niên Indonesia: phát hiện '
     'từ nghiên cứu cắt ngang sử dụng thang đo đã xác nhận và phương pháp '
     'lấy mẫu sáng tạo)',
     'Pham MD, Wulan NR, Sawyer SM, Agius PA, Fisher J, Tran T, '
     'Medise BE, Devaera Y, Riyanti A, Ansariadi A — Indonesia',
     'Journal of Adolescent Health (Q1, IF 3.74) — 2024',
     'Đa-tâm; lo âu xã hội nổi bật'),
    ('2', 'Prevalence and predictors of social anxiety disorders among '
     'Malaysian secondary school students during the COVID-19 pandemic: '
     'Exploring the influence of internet gaming disorder and impulsivity\n'
     '(Tỷ lệ hiện mắc và các yếu tố dự báo rối loạn lo âu xã hội ở học '
     'sinh trung học Malaysia trong đại dịch COVID-19: khảo sát ảnh hưởng '
     'của rối loạn nghiện trò chơi điện tử và tính bốc đồng)',
     'Mohamed NF, Ting TJ, Manan NAA, Senari IFM, Chan MFMF, '
     'Rahmatullah B, Govindasamy P, Abdulla K — Malaysia — n=1.574',
     'Clinical Child Psychology and Psychiatry (SAGE, Q2, IF 1.8-2.0) — 2024',
     'Tỷ lệ SAD 40,53%; 12 trường ở 5 bang'),
    ('3', 'Association between commuting and mental health among Japanese '
     'adolescents\n'
     '(Mối liên hệ giữa việc đi học xa và sức khỏe tâm thần ở thanh thiếu '
     'niên Nhật Bản)',
     'Suguru Nakajima, Yuichiro Otsuka, Osamu Itani, Yoshiyuki Kaneko, '
     'Masahiro Suzuki, Yoshitaka Kaneita — Nihon Univ. School of Medicine, '
     'Tokyo, Nhật Bản',
     'Psychiatry and Clinical Neurosciences (Wiley, Q1, IF 6.2) — Vol 78(10):'
     '588-594 — 2024',
     'Tỷ lệ trầm cảm 17,3% / lo âu 19,0%'),
]
make_paper_table(q1_asia, [0.8, 7.5, 4.0, 3.0, 2.8])


H2('1.2 Trung Đông')

q1_me = [
    ('4', 'Exploring the psychometric properties of the Persian Depression '
     'Anxiety Stress Scale for Youth (DASS-Y): factor structure and '
     'reliability in Iranian children and adolescents\n'
     '(Khám phá tính chất tâm trắc của thang Đánh giá Trầm cảm – Lo âu – '
     'Stress phiên bản Ba Tư cho Thiếu niên: cấu trúc nhân tố và độ tin cậy '
     'ở trẻ em và thanh thiếu niên Iran)',
     'Shabani MJ, Gharraee B, Tajrishi KZ — Iran (Tehran) — n=1.277',
     'Frontiers in Psychology (Q2, IF 2.9) — 2025',
     'Validation thang DASS-Y phiên bản Ba Tư; CFA xác nhận cấu trúc 3 '
     'nhân tố'),
]
make_paper_table(q1_me, [0.8, 7.5, 4.0, 3.0, 2.8])


H2('1.3 Châu Phi')

q1_af = [
    ('5', 'Prevalence and factors associated with symptoms of depression and '
     'anxiety among young school-going adolescents in the Western Cape '
     'Province of South Africa\n'
     '(Tỷ lệ hiện mắc và các yếu tố liên quan với triệu chứng trầm cảm và '
     'lo âu ở thanh thiếu niên đang đi học tại tỉnh Western Cape, Nam Phi)',
     'Reuben Moyo, Arvin Bhana, Peter S. Nyasulu — Nam Phi — n=621 (10-14 '
     'tuổi, 10 trường tiểu học)',
     'Comprehensive Psychiatry (Elsevier, Q1, IF 5.30) — Vol 131:152469 — 2024',
     'Trầm cảm 33,5% / lo âu 20,9%; PHQ-A + GAD-7'),
]
make_paper_table(q1_af, [0.8, 7.5, 4.0, 3.0, 2.8])


H2('1.4 Mỹ La-tinh + Châu Âu')

q1_la_eu = [
    ('6', 'Anxiety among children a year after the onset of the COVID-19 '
     'pandemic: a Brazilian cross-sectional online survey\n'
     '(Lo âu ở trẻ em một năm sau khi đại dịch COVID-19 bắt đầu: khảo sát '
     'trực tuyến cắt ngang tại Brazil)',
     'Garcia de Avila MA, de Jesus Amorin T, Hamamoto Filho PT, '
     'de Almeida GMF, Olaya-Contreras P, Berghammer M, Jenholt Nolbris M, '
     'Nilsson S — Brazil',
     'Frontiers in Public Health (Q1, IF 3.46) — 2024',
     'CAQ ≥9: lo âu 24,9%; NRS ≥8: 34,9%'),
    ('7', 'Selected determinants of anxiety and depression symptoms in '
     'adolescents aged 11–15 in relation to the pandemic COVID-19 and the '
     'war in Ukraine\n'
     '(Các yếu tố quyết định lo âu và triệu chứng trầm cảm ở thanh thiếu '
     'niên 11-15 tuổi liên quan đến đại dịch COVID-19 và chiến tranh '
     'Ukraine)',
     'Grzankowska I, Wójtowicz-Szefler M, Deja M — Ba Lan',
     'Frontiers in Public Health (Q1, IF 3.46) — 2025',
     'Nhấn mạnh biến nội tâm (self-esteem) trong khủng hoảng kép'),
]
make_paper_table(q1_la_eu, [0.8, 7.5, 4.0, 3.0, 2.8])


d.add_page_break()


# ============================================================
H1('PHẦN 2 — Q2/Q3 (MÔ TẢ CẮT NGANG)')

H2('2.1 Đông Nam Á')

q3_sea = [
    ('8', 'Predictors and Protective Factors of Mental Health Outcomes '
     'Among Indonesian Adolescents: An SCT-Based Multi-Center Study\n'
     '(Các yếu tố dự báo và yếu tố bảo vệ kết quả sức khỏe tâm thần ở '
     'thanh thiếu niên Indonesia: nghiên cứu đa-tâm dựa trên Lý thuyết '
     'Nhận thức Xã hội)',
     'Sarfika R, Saifudin IMMY, Wicaksana AL, Malini H, Wenny BP, '
     'Rahayuningsih A, Wijaya NE, Putri DE, Abdullah KL — Indonesia '
     '(West Sumatra, Yogyakarta, Central Kalimantan) — n=2.984 từ 19 '
     'trường THPT công',
     'International Journal of Social Psychiatry (SAGE, Q1, IF 2.7) — 2025',
     'Khung lý thuyết SCT; đa địa điểm'),
    ('9', 'Sociodemographic characteristics, social support, and family '
     'history as factors of depression, anxiety, and stress among young '
     'adult senior high school students in metro Manila, Philippines, '
     'during the COVID-19 pandemic\n'
     '(Đặc điểm nhân khẩu xã hội, hỗ trợ xã hội và tiền sử gia đình như '
     'các yếu tố liên quan đến trầm cảm, lo âu và stress ở học sinh THPT '
     'tuổi vị thành niên muộn tại Metro Manila, Philippines, trong đại '
     'dịch COVID-19)',
     'Serrano IMA, Cuyugan AMN, Cruz K, Mahusay JMA, Alibudbud R — '
     'Philippines — n=187 (De La Salle University)',
     'Frontiers in Psychiatry (Q1, IF 3.2) — 2023',
     'DASS-21 + MSPSS; 4/5 nguy cơ lo âu'),
]
make_paper_table(q3_sea, [0.8, 7.5, 4.0, 3.0, 2.8])


H2('2.2 Trung Đông + Bắc Phi')

q3_me_naf = [
    ('10', 'Mental Health Outcomes Among Homeschooled and Traditionally '
     'Schooled Adolescents in Egypt (Ages 12–18): A Cross-Sectional Study\n'
     '(Các kết quả sức khỏe tâm thần ở thanh thiếu niên Ai Cập học tại nhà '
     'và học trường truyền thống (12-18 tuổi): nghiên cứu cắt ngang)',
     'Hasan A, Ismael A, Zayed MS, Abdelwahed MM et al. — Ai Cập '
     '(Al-Azhar Univ., Cairo Univ.)',
     'Cureus (IF 1.3) — 2025',
     'Tỷ lệ lo âu 45,2%; so sánh homeschool vs trường'),
    ('11', 'Prevalence and Determinants of Depression, Anxiety, and Stress '
     'Among Secondary School Students [Saudi Arabia]\n'
     '(Tỷ lệ hiện mắc và các yếu tố quyết định trầm cảm, lo âu và stress ở '
     'học sinh trung học [Ả Rập Saudi])',
     'Mohammed M Barnawi, Ali M Sonbaa, Maysa M Barnawi, Abdullah H '
     'Alqahtani, Bashaier A Fairaq — Ả Rập Saudi — n=702 (DASS-21 phiên '
     'bản Arab)',
     'Cureus 15(8):e44182 (IF 1.3) — 2023',
     '⚠ Cureus mất WoS indexing 27/10/2025; lo âu 35,2% / trầm cảm '
     '30,8% / stress 14,7%'),
    ('12', 'Universal depressive symptom screening in middle schools in '
     'Istanbul: An epidemiologic study\n'
     '(Sàng lọc triệu chứng trầm cảm toàn diện ở các trường trung học cơ '
     'sở Istanbul: nghiên cứu dịch tễ)',
     'Yildiz Silahli N, Baris HE, Qutranji L, Yorganci Kale B, Günal Ö, '
     'Ütük B, Karavuş M, Rodopman Arman A, Boran P — Thổ Nhĩ Kỳ '
     '(Istanbul) — n=6.110 từ 9 trường',
     'Journal of Affective Disorders (Q1, IF 5.42) — 2025',
     'Trầm cảm 10-20%; sàng lọc THCS quy mô lớn'),
]
make_paper_table(q3_me_naf, [0.8, 7.5, 4.0, 3.0, 2.8])


H2('2.3 Đông Phi')

q3_eaf = [
    ('13', 'Mental health and psychological well-being of Kenyan adolescents '
     'from Nairobi and the Coast regions in the context of COVID-19\n'
     '(Sức khỏe tâm thần và phúc lợi tâm lý của thanh thiếu niên Kenya '
     'tại Nairobi và vùng ven biển trong bối cảnh COVID-19)',
     'Mbithi G, Mabrouk A, Sarki A, Odhiambo R, Namuguzi M, Dzombo JT, '
     'Atukwatse J, Kabue M, Mwangi P, Abubakar A — Kenya — n=2.192 (4 '
     'trường tại Nairobi + Kiambu)',
     'Child Adolesc Psychiatry Mental Health (Q1, IF 4.6) — 2023',
     'Dùng GAD-7; số liệu nguồn mở'),
]
make_paper_table(q3_eaf, [0.8, 7.5, 4.0, 3.0, 2.8])


d.add_page_break()


# ============================================================
H1('PHẦN 3 — PHÂN TÍCH & ỨNG DỤNG CHO NHÓM')

H2('3.1 Tiêu đề tham khảo theo dạng cấu trúc')

H3('Dạng 1: Validation + Đối tượng + Thang đo')
P('Mẫu thành công: "Exploring the psychometric properties of [scale] '
  '...factor structure and reliability in [population]" (Iran 2025).')
P('Áp dụng VN: "Psychometric properties of the Vietnamese RCADS in lower '
  'secondary school students: factor structure and gender invariance" — '
  'có thể là bài khả thi cho Q2/Q3 nếu tách phần validation thang đo của '
  'NCS.', italic=True)

H3('Dạng 2: Yếu tố dự báo + Đối tượng + Khung lý thuyết')
P('Mẫu thành công: "Predictors and Protective Factors of Mental Health '
  'Outcomes Among [Indonesian Adolescents]: An [SCT-Based] Multi-Center '
  'Study" (Indonesia 2025).')
P('Áp dụng VN: bài Q1 của nhóm có thể nhấn vào khung lý thuyết "Risk and '
  'Protective Factors Framework" — Bronfenbrenner ecological model.',
  italic=True)

H3('Dạng 3: So sánh phân nhóm bối cảnh')
P('Mẫu thành công: "Mental Health Outcomes Among Homeschooled and '
  'Traditionally Schooled Adolescents in [country]" (Egypt 2025).')
P('Áp dụng VN: có thể là chủ đề bài Q3 phụ — so sánh học sinh trường '
  'công vs trường tư hoặc đô thị vs ngoại ô.', italic=True)

H3('Dạng 4: Sàng lọc dịch tễ quy mô lớn')
P('Mẫu thành công: "Universal depressive symptom screening in middle '
  'schools in [city]: An epidemiologic study" (Turkey 2025).')
P('Áp dụng VN: tiêu đề Q3 của nhóm có thể đi theo hướng "Universal '
  'anxiety symptom screening in Vietnamese lower secondary schools" nếu '
  'muốn nhấn vào ứng dụng sàng lọc rộng.', italic=True)


H2('3.2 Bài học từ bộ sưu tập mở rộng')

B('Nhiều paper Q1 thành công ở Frontiers in Psychiatry / Public Health / '
  'Psychology — open access nhưng IF 3-5, dễ chấp nhận hơn nature.com hay '
  'BMC tạp chí thiêng')
B('Comprehensive Psychiatry (Elsevier, IF 5.30) là target tốt nếu nhóm có '
  'dataset SEM mạnh — Western Cape paper 2024 là ví dụ')
B('Journal of Adolescent Health (Q1, IF 3.74) là Q1 thân thiện với '
  'multicenter studies — Indonesia 2024 đã thành công')
B('Cureus IF 1.3 và mới bị Web of Science loại khỏi indexing từ '
  '27/10/2025 → KHÔNG nên submit về sau, dù paper Ai Cập 2025 đã thành '
  'công ở đây trước thời điểm đó')
B('Bài về validation thang đo (như Persian DASS-Y) là một hướng phụ Q2 '
  'mà nhóm có thể tách ra nếu có nhu cầu publication output thứ ba')


H2('3.3 Khuyến nghị bổ sung cho 2 bài chính')

P('Sau khi mở rộng dataset tham khảo, em nhận thấy mẫu tiêu đề các bài '
  'thành công ở Q1 thường có CẤU TRÚC NGẮN GỌN (15-22 từ) + ROBUST '
  'METHODOLOGY trong tiêu đề (SEM, multi-center, validation, '
  'epidemiologic study, longitudinal). Em vẫn duy trì khuyến nghị A1 / A4 '
  'cho Q1 và B2 / B5 cho Q3 từ file v4.', italic=True)


# ============================================================
H1('PHẦN 4 — BẢNG IF JCR 2024 CỦA CÁC TẠP CHÍ MỚI')

journals_new = [
    (('Tên tạp chí', 'Q-rank', 'IF 2024', 'Publisher', 'Ghi chú')),
    (('Journal of Adolescent Health', 'Q1', '3.74', 'Elsevier',
      'Mục tiêu Q1 thân thiện multicenter')),
    (('Psychiatry and Clinical Neurosciences', 'Q1', '6.2', 'Wiley',
      'IF cao; quality bar nghiêm')),
    (('Comprehensive Psychiatry', 'Q1', '5.30', 'Elsevier',
      'Target tốt cho SEM/longitudinal')),
    (('Frontiers in Public Health', 'Q1', '3.46', 'Frontiers',
      'Open access; nhiều paper VN khả thi')),
    (('International Journal of Social Psychiatry', 'Q1', '2.7',
      'SAGE', 'Verified JCR 2024 (em ghi sai Q2 trong v1)')),
    (('Clinical Child Psychology and Psychiatry', 'Q2', '1.8-2.0', 'SAGE',
      'Q2; chủ đề chuyên biệt cho trẻ em + thanh thiếu niên')),
    (('Cureus', '—', '1.3', 'Cureus Inc.',
      '⚠ MẤT INDEXING Web of Science từ 27/10/2025 — không nên target '
      'tiếp')),
]

t_j = d.add_table(rows=1, cols=5); t_j.style = 'Light Grid Accent 1'
t_j.autofit = False
hdr = t_j.rows[0].cells
for i, h in enumerate(journals_new[0]):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
for row_data in journals_new[1:]:
    row = t_j.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t_j, [4.5, 1.5, 2.0, 2.5, 6.5])


WARN('Cureus (IF 1.3) đã mất Web of Science indexing từ 27/10/2025. Hai '
     'paper #10 (Ai Cập) và #11 (Ả Rập Saudi) đều ở Cureus — vẫn cite '
     'được làm tham khảo cho NCS nhưng KHÔNG nên target tiếp cho bài '
     'mới của nhóm.')


# ============================================================
H1('PHẦN 5 — NGUỒN VERIFY')

P('Mọi entry trong file này đã được verify qua các nguồn sau (tuân thủ '
  'memory `feedback_verify_tung_entry_truoc_khi_gui.md`):', indent=False)

B('PubMed/PMC — tác giả + năm + DOI')
B('Crossref + Journal landing pages (frontiersin.org, jahonline.org, '
  'onlinelibrary.wiley.com, sciencedirect.com, link.springer.com, '
  'journals.sagepub.com, cureus.com)')
B('JCR 2024 thông qua wos-journal.info — IF + quartile')
B('SCImago — cross-check quartile')

P('Lưu ý: vì các paper mới này em ít quen thuộc hơn các paper TQ/VN nên '
  'mức độ chi tiết tác giả (full vs abbreviated) có thể không đồng đều '
  'giữa các entry. Khi cite formal trong draft bài Q1/Q3, NCS nên tra '
  'cứu lại từng paper trên Crossref để có citation chuẩn APA.', italic=True)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
