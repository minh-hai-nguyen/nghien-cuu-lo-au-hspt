# -*- coding: utf-8 -*-
"""
Round 5 — trim Bài 1 Tóm tắt VN to push ALL ≤6200.
"""
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
B1 = ROOT / "bai-bao-khgdvn/Bai1_YTNC_HSTHCS_v4_14052026.docx"


def set_run_format(run, font_name="Times New Roman", size=11, bold=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


# Trimmed Tóm tắt VN (208 từ instead of 243)
TOM_TAT_NEW = (
    "Rối loạn lo âu là vấn đề sức khỏe tâm thần phổ biến nhất ở vị thành niên trên toàn cầu, "
    "với xu hướng gia tăng trong thập niên qua. Tại Việt Nam, kết quả Điều tra Sức khỏe tâm thần "
    "vị thành niên quốc gia (V-NAMHS) ghi nhận tỷ lệ chẩn đoán DSM-5 đầy đủ cho nhóm lo âu ở mức "
    "2,3% và tỷ lệ vấn đề sức khỏe tâm thần chung khoảng 21,7%, song dữ liệu về yếu tố nguy cơ "
    "ở lứa tuổi trung học cơ sở (11–15 tuổi) vẫn còn rời rạc. Bài viết tổng quan tự sự các "
    "nghiên cứu công bố 2015–2026 nhằm hệ thống hóa năm nhóm yếu tố nguy cơ chính: áp lực học "
    "tập, nghiện điện thoại thông minh, bắt nạt thể chất, bắt nạt bằng lời và bắt nạt mạng, và "
    "lòng tự trọng thấp. Kết quả cho thấy mỗi nhóm đều có bằng chứng quốc tế mạnh và đã được "
    "ghi nhận trong các công trình tại Việt Nam ở quy mô khu vực. Tuy nhiên vẫn còn nhiều khoảng "
    "trống, đặc biệt về dữ liệu nghiên cứu dọc, học sinh dân tộc thiểu số và mô hình nguy cơ kết "
    "hợp. Bài viết đề xuất bốn hướng nghiên cứu ưu tiên cho giai đoạn tới."
)


def main():
    print(f"Trim Tóm tắt VN: {len(TOM_TAT_NEW.split())} từ")
    doc = Document(B1)
    for p in doc.paragraphs:
        if p.text.startswith("Tóm tắt:"):
            if p.runs:
                p.runs[0].text = "Tóm tắt: "
                for r in p.runs[1:]:
                    r.text = ""
                new_run = p.add_run(TOM_TAT_NEW)
                set_run_format(new_run, size=11)
                print(f"  ✓ Replaced Tóm tắt VN")
            break

    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)
    doc.save(B1)
    print(f"  Saved: {B1}")


if __name__ == "__main__":
    main()
