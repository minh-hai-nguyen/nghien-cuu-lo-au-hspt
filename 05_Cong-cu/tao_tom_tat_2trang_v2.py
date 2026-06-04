# -*- coding: utf-8 -*-
"""
Tạo:
1. Bảng tổng hợp 2 trang (chi tiết cột Phát hiện nổi bật 3-5 câu)
2. 11 tóm tắt 2 trang riêng cho mỗi bài
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w * 567)))
    we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def new_doc(narrow=False):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.3
    for s in doc.sections:
        s.top_margin = Cm(2 if not narrow else 1.5)
        s.bottom_margin = Cm(2 if not narrow else 1.5)
        s.left_margin = Cm(2.5 if not narrow else 2)
        s.right_margin = Cm(2 if not narrow else 1.5)
    return doc

def add_h(doc, text, level=1, size=14):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(size)
        r.font.color.rgb = RGBColor(0, 0, 0)

def add_p(doc, text, bold=False, italic=False, size=11, color=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    if color: r.font.color.rgb = color
    return p

def add_info_tbl(doc, data, widths=(3.5, 12.5)):
    t = doc.add_table(rows=len(data), cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        set_w(row.cells[0], widths[0])
        set_w(row.cells[1], widths[1])
    for i, (k, v) in enumerate(data):
        c0 = t.rows[i].cells[0]
        c0.text = k
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade(c0, 'D9E2F3')
        c1 = t.rows[i].cells[1]
        c1.text = v
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)

# ==============================================================
# DATA for all 11 papers
# ==============================================================
papers = [
    {
        'num': '01', 'author': 'Jenkins và cs.', 'year': '2023',
        'title_vi': 'Trầm cảm và lo âu ở học sinh THCS đa sắc tộc: Tuổi, giới tính và môi trường văn hóa xã hội',
        'title_en': 'Depression and anxiety among multiethnic middle school students: Age, gender, and sociocultural environment',
        'journal': 'International Journal of Social Psychiatry, 69(3), 784\u2013794',
        'doi': '10.1177/00207640221140282',
        'country': 'Hoa Kỳ (San Diego, California)',
        'sample': '75 học sinh THCS đa sắc tộc (tuổi TB 11,2)',
        'tools': 'PHQ-9A (trầm cảm), GAD-10 (lo âu), phỏng vấn dân tộc học bán cấu trúc',
        'design': 'Nghiên cứu hỗn hợp (mixed methods): cắt ngang + dân tộc học, 2018\u20132021',
        'anxiety': '50,6% (GAD-10, bất kỳ mức)',
        'depression': '44% (PHQ-9A, bất kỳ mức)',
        'findings': 'Phương pháp hỗn hợp kết hợp sàng lọc định lượng (PHQ-9A, GAD-10) với phỏng vấn dân tộc học mở \u2014 cho phép phát hiện các yếu tố văn hóa xã hội mà thang đo không nắm bắt được. Nữ sinh có điểm trầm cảm cao hơn nam (p = 0,002, Mann-Whitney) và lo âu cao hơn (p = 0,016). Phỏng vấn phát hiện ba bối cảnh gây căng thẳng chính: bạo lực trên cơ sở giới (trong trường và gia đình), lo lắng về COVID-19 (bệnh tật, tử vong, mất việc, khó khăn kinh tế), và sự gia tăng phân biệt chủng tộc. Học sinh trải nghiệm những vấn đề này như rào cản đối với sự tham gia giáo dục.',
        'critique': 'Cỡ mẫu rất nhỏ (n=75), chọn mẫu thuận tiện tại 1 trường duy nhất. GAD-10 là phiên bản sửa đổi thiếu dữ liệu xác thực (validation). Thiết kế kéo dài 2018\u20132021 trùng COVID-19 nhưng không kiểm soát hiệu ứng thời điểm. Tuy nhiên, phương pháp hỗn hợp là điểm sáng \u2014 ít NC nào kết hợp dân tộc học với sàng lọc.',
        'gap': 'Cần mở rộng cỡ mẫu. Phát triển công cụ sàng lọc nhạy cảm văn hóa. Nghiên cứu dọc theo dõi từ đầu vị thành niên. Đánh giá can thiệp giải quyết yếu tố văn hóa xã hội.',
        'quality': '\u2b50\u2b50\u2b50 Trung bình-Khá',
    },
    {
        'num': '02', 'author': 'Saikia và cs.', 'year': '2023',
        'title_vi': 'Các bệnh lý SKTT và yếu tố liên quan ở thanh thiếu niên tại Kamrup, Assam',
        'title_en': 'Mental Health Morbidities and their Correlates among the Adolescents in Kamrup (Metro), Assam',
        'journal': 'Indian Journal of Community Medicine, 48(6), 835\u2013840',
        'doi': '10.4103/ijcm.ijcm_574_23',
        'country': '\u1ea4n \u0110\u1ed9 (\u0110\u00f4ng B\u1eafc, Assam)',
        'sample': '360 h\u1ecdc sinh 14\u201319 tu\u1ed5i t\u1ea1i 4 tr\u01b0\u1eddng',
        'tools': 'DASS-21 (Depression Anxiety Stress Scale \u2014 21 m\u1ee5c)',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang t\u1ea1i tr\u01b0\u1eddng h\u1ecdc, th\u00e1ng 4/2019\u20136/2020',
        'anxiety': '24,4% (DASS-21)',
        'depression': '22,2%',
        'findings': 'Phát hiện đáng chú ý nhất: nam giới có tỷ lệ lo âu cao hơn nữ (30,0% vs 18,9%, P=0,049) \u2014 trái ngược hoàn toàn với y văn quốc tế và hầu hết các nghiên cứu khác. Đây là nghiên cứu tiên phong tại vùng Đông Bắc Ấn Độ, nơi có đặc điểm văn hóa bộ lạc khác biệt. Các yếu tố nguy cơ: cha mẹ đơn thân, sử dụng rượu, nghiện game, kết quả học tập kém. Tỷ lệ SKTT tổng thể (24,4% lo âu) thấp hơn nhiều so với nghiên cứu tại miền Nam Ấn Độ \u2014 gợi ý sự khác biệt vùng miền đáng kể.',
        'critique': 'DASS-21 phiên bản tiếng Assam thiếu báo cáo Cronbach alpha. Chỉ dùng chi-square (đơn biến), không hồi quy đa biến. Thu thập dữ liệu 4/2019\u20136/2020 trùng COVID-19 nhưng không đề cập tác động. Kết quả nam > nữ cần giải thích sâu hơn về bối cảnh văn hóa bộ lạc Đông Bắc.',
        'gap': 'Nghiên cứu quy mô lớn hơn ở Đông Bắc Ấn Độ. Khảo sát yếu tố văn hóa bộ lạc đặc thù. So sánh thành thị-nông thôn. Giải thích tại sao nam > nữ.',
        'quality': '\u2b50\u2b50\u2b50 Trung bình',
    },
    {
        'num': '03', 'author': 'Mandaknalli & Malusare', 'year': '2021',
        'title_vi': 'Nghiên cứu cắt ngang về tỷ lệ lo âu ở khu vực trường học thuộc đô thị tự trị',
        'title_en': 'A cross-sectional study on the prevalence of anxiety among municipality school area',
        'journal': 'MedPulse Int. J. Psychology, 18(3), 19\u201322',
        'doi': '',
        'country': '\u1ea4n \u0110\u1ed9',
        'sample': '450 h\u1ecdc sinh',
        'tools': 'B\u1ea3ng c\u00e2u h\u1ecfi lo \u00e2u (kh\u00f4ng n\u00eau t\u00ean c\u1ee5 th\u1ec3)',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang',
        'anxiety': 'Nh\u1eb9 49,4%, Trung b\u00ecnh 43,3%, N\u1eb7ng 7,3%',
        'depression': 'Kh\u00f4ng \u0111o',
        'findings': 'Gần 100% học sinh có ít nhất mức lo âu nhẹ \u2014 gợi ý điểm cắt quá thấp hoặc công cụ đo quá nhạy. Nữ sinh có tỷ lệ lo âu nặng cao hơn nam (10,9% vs 3,8%). Các yếu tố nguy cơ: thiếu hoạt động thể chất, chất lượng giấc ngủ kém, và hút thuốc lá. Nghiên cứu gợi ý can thiệp tập trung vào thay đổi lối sống (tập thể dục, cải thiện giấc ngủ) có thể giảm lo âu.',
        'critique': 'Thiếu nghiêm trọng thông tin phương pháp: công cụ không nêu tên cụ thể, không có Cronbach alpha. Gần 100% có lo âu là bất thường. Tạp chí không có IF, không PubMed. Không có chấp thuận đạo đức nghiên cứu.',
        'gap': 'Tái thực hiện với công cụ đo chuẩn hóa (DASS-21 hoặc GAD-7). Nghiên cứu vai trò hoạt động thể chất trong can thiệp lo âu tại trường.',
        'quality': '\u2b50\u2b50 Trung bình-Thấp',
    },
    {
        'num': '04', 'author': 'NSCH', 'year': '2020',
        'title_vi': 'Khảo sát Quốc gia về Sức khỏe Trẻ em (NSCH) 2020',
        'title_en': 'National Survey of Children\u2019s Health (NSCH) 2020',
        'journal': 'HRSA/Census Bureau Report',
        'doi': '',
        'country': 'Hoa Kỳ (quốc gia)',
        'sample': '55.162 trẻ em và thanh thiếu niên (18.397 TN 12\u201317)',
        'tools': 'Kh\u1ea3o s\u00e1t qu\u1ed1c gia chu\u1ea9n h\u00f3a (d\u1ef1a tr\u00ean ch\u1ea9n \u0111o\u00e1n y t\u1ebf)',
        'design': 'Kh\u1ea3o s\u00e1t qu\u1ed1c gia h\u00e0ng n\u0103m, d\u1eef li\u1ec7u 2016\u20132023',
        'anxiety': '16,1% (2020, ch\u1ea9n \u0111o\u00e1n); t\u0103ng 61% t\u1eeb 2016',
        'depression': '8,4% (2020, ch\u1ea9n \u0111o\u00e1n)',
        'findings': 'Nguồn dữ liệu quốc gia uy tín nhất về SKTT trẻ em Hoa Kỳ. Lo âu tăng 61% trong 7 năm (2016\u20132023) \u2014 xu hướng tăng rõ ràng và liên tục. Nữ giới có tỷ lệ cao hơn nam (20,1% vs 12,3%). Thanh thiếu niên bị lo âu có nguy cơ khó kết bạn gấp 10 lần và nghỉ học gấp 5 lần. 61% gia đình gặp khó khăn khi tiếp cận dịch vụ điều trị SKTT. Tỷ lệ dựa trên chẩn đoán y tế (không phải sàng lọc) nên thấp hơn nhưng chính xác hơn.',
        'critique': 'Dữ liệu từ cha mẹ báo cáo, không trực tiếp từ trẻ \u2014 có thể underestimate (so với Alharbi sàng lọc: 74% vs NSCH 16,1%). Dựa trên chẩn đoán y tế \u2014 phụ thuộc tiếp cận dịch vụ. Xu hướng tăng 61% có thể phản ánh giảm kỳ thị và tăng nhận thức hơn là tăng bệnh thực sự.',
        'gap': 'Phân tích nguyên nhân gia tăng 61%. Nghiên cứu giải pháp giảm rào cản tiếp cận điều trị. Đánh giá tác động dài hạn COVID-19.',
        'quality': '\u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc',
    },
    {
        'num': '05', 'author': 'Alharbi và cs.', 'year': '2019',
        'title_vi': 'Trầm cảm và lo âu ở HS THPT tại vùng Qassim, Ả Rập Saudi',
        'title_en': 'Depression and anxiety among high school student at Qassim Region',
        'journal': 'J Family Med Prim Care, 8(2), 504\u2013510',
        'doi': '10.4103/jfmpc.jfmpc_383_18',
        'country': '\u1ea2 R\u1eadp Saudi (v\u00f9ng Qassim)',
        'sample': '1.245 h\u1ecdc sinh THPT',
        'tools': 'PHQ-9 (tr\u1ea7m c\u1ea3m), GAD-7 (lo \u00e2u)',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang, ch\u1ecdn m\u1eabu ng\u1eabu nhi\u00ean nhi\u1ec1u giai \u0111o\u1ea1n',
        'anxiety': '63,5% (GAD-7, b\u1ea5t k\u1ef3 m\u1ee9c)',
        'depression': '74% (PHQ-9, b\u1ea5t k\u1ef3 m\u1ee9c)',
        'findings': 'Tỷ lệ cực cao: 74% trầm cảm và 63,5% lo âu \u2014 tuy nhiên dùng ngưỡng PHQ-9 \u2265 5 (bao gồm cả nhẹ), nên cao hơn nhiều so với ngưỡng lâm sàng. Là bài đầu tiên tại Ả Rập Saudi sử dụng PHQ-9/GAD-7 \u2014 mở ra hướng nghiên cứu mới cho khu vực. Nữ giới có tỷ lệ cao hơn nam (P<0,001). So sánh: Chen 2023 dùng cùng công cụ nhưng ngưỡng cao hơn cho 23% trầm cảm, cho thấy tầm quan trọng của ngưỡng cắt.',
        'critique': 'Ngưỡng PHQ-9 \u2265 5 quá thấp (bao gồm nhẹ). Chỉ dùng chi-square, không hồi quy đa biến. PHQ-9/GAD-7 phiên bản phương Tây chưa xác thực cho bối cảnh Ả Rập Saudi. Thiên lệch địa lý (Buraidah chiếm 28,8% mẫu).',
        'gap': 'Nghiên cứu dọc xu hướng theo thời gian. Khảo sát vai trò mạng xã hội và áp lực văn hóa Ả Rập. Chương trình SKTT tại trường ở Ả Rập Saudi.',
        'quality': '\u2b50\u2b50\u2b50\u2b50 Khá-Tốt',
    },
    {
        'num': '06', 'author': 'Nakie và cs.', 'year': '2022',
        'title_vi': 'Tỷ lệ và yếu tố liên quan của trầm cảm, lo âu, căng thẳng ở HS THPT Tây Bắc Ethiopia',
        'title_en': 'Prevalence and associated factors of depression, anxiety, and stress among high school students in Northwest Ethiopia, 2021',
        'journal': 'BMC Psychiatry, 22:739',
        'doi': '10.1186/s12888-022-04389-x',
        'country': 'Ethiopia (T\u00e2y B\u1eafc)',
        'sample': '849 h\u1ecdc sinh THPT (15\u201325 tu\u1ed5i)',
        'tools': 'DASS-21',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang, ch\u1ecdn m\u1eabu ng\u1eabu nhi\u00ean \u0111\u01a1n gi\u1ea3n',
        'anxiety': '66,7% (DASS-21)',
        'depression': '41,4%',
        'findings': 'Tỷ lệ lo âu 66,7% và căng thẳng 52,2% rất cao \u2014 có thể liên quan đến bối cảnh xung đột Ethiopia 2020\u20132022. Là nghiên cứu đầu tiên tại châu Phi đánh giá đồng thời cả 3 tình trạng (trầm cảm, lo âu, căng thẳng). Yếu tố nguy cơ mạnh nhất: nhai lá khat (AOR=5,6), hút thuốc (AOR=4,8), và bệnh mãn tính (AOR=2,1). Đăng trên BMC Psychiatry (Q1, IF \u2248 4,4) \u2014 uy tín cao.',
        'critique': 'Phạm vi tuổi rộng (15\u201325) bao gồm cả người trưởng thành. AOR hút thuốc CI trên = 7304 \u2014 có thể lỗi in ấn. DASS-21 tiếng Amharic thiếu thông tin xác thực văn hóa. Thiết kế cắt ngang không phân biệt nguyên nhân/hệ quả sử dụng chất.',
        'gap': 'Mở rộng sang khu vực khác của Ethiopia và châu Phi. Can thiệp kép: giảm sử dụng chất + hỗ trợ SKTT. Nghiên cứu dọc đánh giá diễn tiến tự nhiên.',
        'quality': '\u2b50\u2b50\u2b50\u2b50 Tốt',
    },
    {
        'num': '07', 'author': 'Chen và cs.', 'year': '2023',
        'title_vi': 'Tỷ lệ và yếu tố liên quan của triệu chứng trầm cảm và lo âu ở HS trung học Trung Quốc',
        'title_en': 'Prevalence and associated factors of depressive and anxiety symptoms among Chinese secondary school students',
        'journal': 'BMC Psychiatry, 23:580',
        'doi': '10.1186/s12888-023-05039-4',
        'country': 'Trung Qu\u1ed1c (mi\u1ec1n T\u00e2y)',
        'sample': '63.205 h\u1ecdc sinh trung h\u1ecdc',
        'tools': 'PHQ-9, GAD-7, PSQI (gi\u1ea5c ng\u1ee7), IGDS9-SF (nghi\u1ec7n game)',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang, l\u1ea5y m\u1eabu c\u1ee5m, c\u00f3 tr\u1ecdng s\u1ed1',
        'anxiety': '13,9% (GAD-7)',
        'depression': '23,0% (PHQ-9)',
        'findings': 'Cỡ mẫu rất lớn (n=63.205), là nghiên cứu đầu tiên tại miền Tây Trung Quốc. Sử dụng nhiều công cụ xác thực: PHQ-9, GAD-7, PSQI (giấc ngủ), IGDS9-SF (nghiện game). Phát hiện 3 yếu tố nguy cơ chính: bắt nạt tại trường, chất lượng giấc ngủ kém, và nghiện game. Gia đình không hạt nhân (đơn thân, tái hôn) cũng là yếu tố nguy cơ đáng kể. Đăng BMC Psychiatry Q1.',
        'critique': 'CI khá rộng dù n lớn \u2014 do hiệu ứng thiết kế lấy mẫu cụm. PHQ-9/GAD-7 là sàng lọc, không phải chẩn đoán. Thiếu mô hình đa tầng (multilevel model) phù hợp cấu trúc cụm. Giấc ngủ kém/chơi game có thể là triệu chứng, không phải nguyên nhân.',
        'gap': 'So sánh Đông-Tây Trung Quốc. Can thiệp nhắm vào bắt nạt + giấc ngủ + game. Nghiên cứu dọc đánh giá quỹ đạo triệu chứng.',
        'quality': '\u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc',
    },
    {
        'num': '08', 'author': 'Wen và cs.', 'year': '2020',
        'title_vi': 'Phân tích hồ sơ tiềm ẩn về lo âu ở HS THCS tại vùng nông thôn kém phát triển của Trung Quốc',
        'title_en': 'A Latent Profile Analysis of Anxiety among Junior High School Students in Less Developed Rural Regions of China',
        'journal': 'Int J Environ Res Public Health, 17(11):4079',
        'doi': '10.3390/ijerph17114079',
        'country': 'Trung Qu\u1ed1c (n\u00f4ng th\u00f4n, 1 t\u1ec9nh)',
        'sample': '900 h\u1ecdc sinh THCS (l\u1edbp 7\u20139)',
        'tools': 'Thang \u0111o lo \u00e2u + Ph\u00e2n t\u00edch h\u1ed3 s\u01a1 ti\u1ec1m \u1ea9n (LPA \u2014 Latent Profile Analysis)',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang, ph\u00e2n t\u00edch LPA b\u1eb1ng Mplus 7.4',
        'anxiety': '24,78% lo \u00e2u n\u1eb7ng (LPA)',
        'depression': 'Kh\u00f4ng \u0111o ri\u00eang',
        'findings': 'Là nghiên cứu đầu tiên sử dụng Phân tích Hồ sơ Tiềm ẩn (LPA) ở HS nông thôn Trung Quốc. LPA xác định 3 nhóm lo âu khác biệt (thấp/trung bình/cao) thay vì chỉ dùng ngưỡng cắt. Nam giới có tỷ lệ lo âu cao hơn nữ \u2014 phù hợp với Saikia (2023) và Xu (2021), trái ngược y văn phương Tây. Hỗ trợ tâm thần tại trường là yếu tố bảo vệ duy nhất có ý nghĩa \u2014 bằng chứng mạnh cho việc đầu tư tư vấn tâm lý trường học.',
        'critique': 'LPA thiếu báo cáo chỉ số phù hợp mô hình (BIC, AIC, entropy). n=900 từ 1 tỉnh \u2014 hạn chế khái quát hóa. Công cụ đo lo âu không nêu tên cụ thể. "Hỗ trợ tại trường" đo tự đánh giá \u2014 có thể thiên lệch nhận thức.',
        'gap': 'Áp dụng LPA ở nhiều vùng nông thôn. Nghiên cứu dọc theo dõi các hồ sơ lo âu. Đánh giá hiệu quả dịch vụ tư vấn tại trường nông thôn.',
        'quality': '\u2b50\u2b50\u2b50\u2b50 Tốt',
    },
    {
        'num': '09', 'author': 'Qiu và cs.', 'year': '2022',
        'title_vi': 'Mối liên hệ giữa phong cách nuôi dạy con và khả năng phục hồi với trầm cảm và lo âu ở HS THCS Trung Quốc',
        'title_en': 'Associations of Parenting Style and Resilience With Depression and Anxiety Symptoms in Chinese Middle School Students',
        'journal': 'Frontiers in Psychology, 13:897339',
        'doi': '10.3389/fpsyg.2022.897339',
        'country': 'Trung Qu\u1ed1c',
        'sample': '2.079 h\u1ecdc sinh THCS (2.879 ban \u0111\u1ea7u, 72,2% ho\u00e0n th\u00e0nh)',
        'tools': 'EMBU (nu\u00f4i d\u1ea1y con), CES-D (tr\u1ea7m c\u1ea3m), SAS (lo \u00e2u), CD-RISC (ph\u1ee5c h\u1ed3i)',
        'design': 'Nghi\u00ean c\u1ee9u c\u1eaft ngang, 4 c\u00f4ng c\u1ee5 \u0111o l\u01b0\u1eddng',
        'anxiety': '13,4% (SAS)',
        'depression': '26,0% (CES-D)',
        'findings': 'Nuôi dạy tiêu cực (trừng phạt, kiểm soát quá mức) tăng nguy cơ trầm cảm (OR=1,82\u20132,01) và lo âu (OR=1,82\u20132,01). Nuôi dạy tích cực (ấm áp, hỗ trợ) giảm nguy cơ mạnh (OR=0,30\u20130,32). Khả năng phục hồi thấp là yếu tố nguy cơ rất mạnh: OR trầm cảm = 6,74. Tuy nhiên, giả thuyết trung tâm của bài (tương tác nuôi dạy \u00d7 phục hồi) KHÔNG có ý nghĩa thống kê \u2014 giảm giá trị đóng góp mới.',
        'critique': 'Giả thuyết chính không được xác nhận. EMBU đo từ quan điểm TTN \u2014 thiên lệch nhận thức tiêu cực ở người trầm cảm. Mẫu 63,4% nam \u2014 mất cân đối giới. CES-D/SAS khác PHQ-9/GAD-7 \u2014 khó so sánh.',
        'gap': 'Nghiên cứu dọc kiểm tra quan hệ nhân quả nuôi dạy \u2192 SKTT. Can thiệp gia đình + tăng cường phục hồi. Nghiên cứu đa văn hóa.',
        'quality': '\u2b50\u2b50\u2b50\u2b50 Tốt',
    },
    {
        'num': '10', 'author': 'Xu và cs.', 'year': '2021',
        'title_vi': 'Tỷ lệ và yếu tố nguy cơ lo âu trong COVID-19: Khảo sát 373.216 HS tại Trung Quốc',
        'title_en': 'Prevalence and risk factors for anxiety symptoms during the outbreak of COVID-19: A large survey among 373,216 junior and senior high school students in China',
        'journal': 'Journal of Affective Disorders, 288, 17\u201322',
        'doi': '10.1016/j.jad.2021.03.080',
        'country': 'Trung Qu\u1ed1c (H\u00e0 Nam \u2014 gi\u00e1p H\u1ed3 B\u1eafc)',
        'sample': '373.216 h\u1ecdc sinh THCS v\u00e0 THPT',
        'tools': 'GAD-7 (lo \u00e2u)',
        'design': 'Kh\u1ea3o s\u00e1t c\u1eaft ngang quy m\u00f4 l\u1edbn, thu th\u1eadp 8 ng\u00e0y (4\u201312/2/2020)',
        'anxiety': '9,89% (GAD-7, \u2265 5)',
        'depression': 'Kh\u00f4ng \u0111o',
        'findings': 'CỠ MẪU LỚN NHẤT TOÀN CẦU về lo âu ở HS (n=373.216). Tỷ lệ lo âu 9,89% \u2014 thấp hơn nhiều NC khác dù trong đỉnh COVID. Nam giới lo âu cao hơn nữ (10,11% vs 9,66%) \u2014 trái y văn (phù hợp Saikia, Wen). Nông thôn cao nhất (12,80%) \u2014 trái ngược kỳ vọng. COVID worry/fear level là yếu tố nguy cơ chính. Đăng J Affective Disorders (Q1, IF \u2248 6,6) \u2014 tạp chí hàng đầu.',
        'critique': 'Thu thập chỉ 8 ngày (4\u201312/2/2020) đỉnh COVID \u2014 đo lo âu tình huống, không phải rối loạn mãn tính. Sức mạnh thống kê quá cao: khác biệt 0,45% (nam vs nữ) đạt P<0,001 mà thiếu ý nghĩa lâm sàng. Chỉ tỉnh Hà Nam \u2014 giáp Hồ Bắc (tâm dịch). Biến "worried/fear level" có thể tautology với lo âu.',
        'gap': 'Nghiên cứu dọc hậu COVID. So sánh quốc tế cùng cỡ mẫu. Hệ thống cảnh báo sớm SKTT trong khủng hoảng. Cơ chế chênh lệch thành thị-nông thôn.',
        'quality': '\u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc',
    },
    {
        'num': '11', 'author': 'Bhardwaj và cs.', 'year': '2020',
        'title_vi': 'Đánh giá trầm cảm, lo âu và căng thẳng ở HS THPT trường công tại Chandigarh, Ấn Độ',
        'title_en': 'A Descriptive study to assess Depression, Anxiety & Stress among higher secondary students of Government schools of Chandigarh, India',
        'journal': 'Journal of IPHA Chandigarh State Branch',
        'doi': '',
        'country': '\u1ea4n \u0110\u1ed9 (Chandigarh)',
        'sample': '288 h\u1ecdc sinh THPT tr\u01b0\u1eddng c\u00f4ng',
        'tools': 'DASS-21',
        'design': 'Nghi\u00ean c\u1ee9u m\u00f4 t\u1ea3 c\u1eaft ngang',
        'anxiety': '73,3% (DASS-21)',
        'depression': '64,9%',
        'findings': 'Tỷ lệ cao nhất trong 11 bài: lo âu 73,3%, trầm cảm 64,9%, căng thẳng 74,7%. Lo âu nặng + cực nặng chiếm 46,8% \u2014 gần một nửa mẫu. Đối tượng là HS trường công lập với hoàn cảnh kinh tế thấp \u2014 gợi ý mối liên hệ mạnh giữa điều kiện kinh tế và SKTT. Tuy nhiên, tỷ lệ cực cao có thể do DASS-21 với ngưỡng thấp trong bối cảnh dân số có nguy cơ cao.',
        'critique': 'Tạp chí địa phương, không PubMed, IF không có. Cỡ mẫu nhỏ (n=288). Thiếu phân tích yếu tố liên quan (không hồi quy). Tỷ lệ cực cao cần xem xét ngưỡng cắt DASS-21.',
        'gap': 'Tái thực hiện với thiết kế phân tích mạnh hơn (hồi quy đa biến). So sánh trường công-tư. Can thiệp SKTT tại trường công. Khảo sát vai trò kinh tế xã hội chi tiết.',
        'quality': '\u2b50\u2b50 Trung bình-Thấp',
    },
]

# ==============================================================
# 1. BẢNG TỔNG HỢP 2 TRANG (chi tiết hơn)
# ==============================================================
doc = new_doc(narrow=True)

title = doc.add_heading('', level=0)
r = title.add_run('BẢNG TÓM TẮT TỔNG HỢP 11 BÀI NGHIÊN CỨU\nLO ÂU VÀ TRẦM CẢM Ở HỌC SINH')
r.font.name = 'Times New Roman'
r.font.size = Pt(13)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

add_p(doc, '6 quốc gia | Cỡ mẫu từ 75 đến 373.216 | 2019\u20132023', italic=True, size=9)

# Table
headers = ['#', 'Tác giả (Năm)', 'Quốc gia\nn', 'Công cụ', 'Lo âu\nTrầm cảm', 'Phát hiện nổi bật (3\u20135 câu)', 'CL']
widths = [0.4, 1.4, 1.2, 0.9, 1.3, 7.8, 0.5]

t = doc.add_table(rows=1+len(papers), cols=len(headers))
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER

for row in t.rows:
    for ci in range(len(headers)):
        set_w(row.cells[ci], widths[ci])

for i, h in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h
    for pp in c.paragraphs:
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in pp.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(6.5)
    shade(c, 'D9E2F3')

for ri, p in enumerate(papers):
    vals = [
        p['num'],
        f"{p['author']}\n({p['year']})",
        f"{p['country'].split('(')[0].strip()}\nn={p['sample'].split()[0]}",
        p['tools'].split('(')[0].strip()[:15],
        f"LA: {p['anxiety']}\nTC: {p['depression']}",
        p['findings'],
        p['quality'].split()[0],
    ]
    for ci, v in enumerate(vals):
        c = t.rows[ri+1].cells[ci]
        c.text = str(v)
        for pp in c.paragraphs:
            for r in pp.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(6.5)

add_p(doc, 'CL = Chất lượng: \u2b50\u2b50 Thấp | \u2b50\u2b50\u2b50 Trung bình | \u2b50\u2b50\u2b50\u2b50 Tốt | \u2b50\u2b50\u2b50\u2b50\u2b50 Xuất sắc. LA = Lo âu, TC = Trầm cảm.', italic=True, size=7)

doc.save('Tóm tắt 2 trang - 11 bài nghiên cứu.docx')

# ==============================================================
# 2. 11 TÓM TẮT RIÊNG (mỗi bài 2 trang)
# ==============================================================
os.makedirs('Tom-tat-tung-bai', exist_ok=True)

for p in papers:
    d = new_doc()

    # Title
    add_h(d, f"Bài {p['num']}: {p['title_vi']}", 1, 13)
    add_p(d, p['title_en'], italic=True, size=10)
    d.add_paragraph()

    # Info table
    add_info_tbl(d, [
        ('Tác giả', f"{p['author']} ({p['year']})"),
        ('Tạp chí', p['journal']),
        ('DOI', p['doi'] if p['doi'] else 'Không có'),
        ('Quốc gia', p['country']),
        ('Mẫu', p['sample']),
        ('Công cụ', p['tools']),
        ('Thiết kế', p['design']),
    ])
    d.add_paragraph()

    # Kết quả chính
    add_h(d, 'KẾT QUẢ CHÍNH', 2, 12)

    # Tỷ lệ
    add_p(d, f"Tỷ lệ lo âu: {p['anxiety']}", bold=True, size=11)
    add_p(d, f"Tỷ lệ trầm cảm: {p['depression']}", bold=True, size=11)
    d.add_paragraph()

    # Phát hiện nổi bật
    add_h(d, 'PHÁT HIỆN NỔI BẬT', 2, 12)
    add_p(d, p['findings'], size=11)
    d.add_paragraph()

    # Phản biện
    add_h(d, 'PHẢN BIỆN', 2, 12)
    pr = d.add_paragraph()
    r = pr.add_run(p['critique'])
    r.font.name = 'Times New Roman'
    r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(255, 0, 0)
    d.add_paragraph()

    # Gap
    add_h(d, 'HƯỚNG NGHIÊN CỨU TIẾP THEO', 2, 12)
    pr2 = d.add_paragraph()
    r2 = pr2.add_run(p['gap'])
    r2.font.name = 'Times New Roman'
    r2.font.size = Pt(11)
    r2.font.color.rgb = RGBColor(0, 0, 200)
    d.add_paragraph()

    # Chất lượng
    add_p(d, f"Đánh giá chất lượng: {p['quality']}", bold=True, size=11)

    fname = f"Tom-tat-tung-bai/{p['num']}_{p['author'].split()[0]}_{p['year']}.docx"
    d.save(fname)

sys.stderr.write('Done: 1 bang tong hop + 11 tom tat rieng\n')
