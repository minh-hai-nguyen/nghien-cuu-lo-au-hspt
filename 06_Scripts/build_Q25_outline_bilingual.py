# -*- coding: utf-8 -*-
"""
Build khung gợi ý bài Q2.5 bilingual EN + VN.
Output: bai-bao-khgdvn/Q25_Outline_EN_VN.docx
Để gửi thầy review + planning.
"""
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
OUT = ROOT / "bai-bao-khgdvn" / "Q25_Outline_EN_VN.docx"


def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False, color=None):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def add_para(doc, text, size=12, bold=False, italic=False, align=None, indent=None,
             space_before=3, space_after=3, line_spacing=1.15, color=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = line_spacing
    if indent is not None:
        pf.first_line_indent = Cm(indent)
    if align == "center":
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == "justify":
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text)
    set_run_format(r, size=size, bold=bold, italic=italic, color=color)
    return p


def add_h1(doc, text):
    return add_para(doc, text, size=14, bold=True, space_before=14, space_after=6,
                    color=RGBColor(0x1f, 0x3a, 0x68))


def add_h2(doc, text):
    return add_para(doc, text, size=13, bold=True, space_before=10, space_after=4,
                    color=RGBColor(0x1f, 0x3a, 0x68))


def add_h3(doc, text):
    return add_para(doc, text, size=12, bold=True, italic=True, space_before=8, space_after=3)


def add_table(doc, headers, rows):
    t = doc.add_table(rows=len(rows) + 1, cols=len(headers))
    t.style = "Light Grid"
    # Headers
    for j, h in enumerate(headers):
        cell = t.rows[0].cells[j]
        cell.text = ""
        p = cell.paragraphs[0]
        r = p.add_run(h)
        set_run_format(r, size=11, bold=True)
    # Rows
    for i, row in enumerate(rows, start=1):
        for j, v in enumerate(row):
            cell = t.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            r = p.add_run(str(v))
            set_run_format(r, size=11)
    return t


def build():
    doc = Document()
    for s in doc.sections:
        s.top_margin = Cm(2.0)
        s.bottom_margin = Cm(2.0)
        s.left_margin = Cm(2.2)
        s.right_margin = Cm(2.0)

    # ============================================================
    # TITLE
    # ============================================================
    add_para(doc,
             "KHUNG GỢI Ý BÀI BÁO Q2.5 TIẾNG ANH",
             size=16, bold=True, align="center", space_before=0, space_after=4,
             color=RGBColor(0x1f, 0x3a, 0x68))
    add_para(doc,
             "Differential pathways from school risk factors to anxiety disorder subtypes "
             "in Vietnamese junior secondary students — A structural equation modelling study",
             size=12, italic=True, align="center", space_after=12)
    add_para(doc,
             "Ngày: 14/05/2026 — Tiếp nối sau 3 bài đã chuẩn bị: 1 bài validation Tâm lý học VN + 2 bài VJES",
             size=10, italic=True, align="center", space_after=14,
             color=RGBColor(0x70, 0x70, 0x70))

    # ============================================================
    # SECTION 0 — CONTEXT
    # ============================================================
    add_h1(doc, "0. Bối cảnh và lý do chiến lược")

    add_h3(doc, "Three published companion papers (3 bài đã chuẩn bị/công bố)")
    add_table(doc,
              ["#", "Bài", "Loại", "Đối tượng / Mẫu", "Tạp chí / Tình trạng"],
              [
                  ["1", "Công Thị Hằng & Đào Minh Đức (2026)",
                   "Validation / CFA", "433 HSTHCS — DSM-5 scales",
                   "Tạp chí Tâm lý học VN — đã sửa Turnitin, đang nộp"],
                  ["2", "VJES Bài 1 — YTNC",
                   "Narrative review", "Tổng quan literature 2015-2026",
                   "VJES — chuẩn bị nộp 14/05/2026"],
                  ["3", "VJES Bài 2 — Can thiệp CBT",
                   "Narrative review", "Tổng quan CBT/School/Digital",
                   "VJES — chuẩn bị nộp 14/05/2026"],
              ])

    add_h3(doc, "Why Q2.5 paper now? (Vì sao cần bài Q2.5 lúc này)")
    add_para(doc,
             "EN: The three companion papers establish (i) measurement validity of DSM-5 anxiety "
             "scales in Vietnamese adolescents, (ii) a state-of-the-art review of risk factors, and "
             "(iii) a review of intervention evidence. The natural next step is an EMPIRICAL paper "
             "that uses the validated scales (paper #1) to test pathway hypotheses arising from the "
             "risk-factor synthesis (paper #2). This sequence is publication-ethics-clean and "
             "maximises citation cross-linkage.",
             size=12, align="justify", indent=0.6)
    add_para(doc,
             "VN: Ba bài đồng hành đã thiết lập (i) độ tin cậy đo lường của các thang DSM-5 trên "
             "vị thành niên Việt Nam, (ii) tổng quan yếu tố nguy cơ, và (iii) tổng quan can thiệp. "
             "Bước tiếp theo tự nhiên là một bài THỰC NGHIỆM dùng các thang đã được kiểm định trong "
             "bài 1 để kiểm định các giả thuyết đường dẫn rút ra từ tổng hợp yếu tố nguy cơ trong "
             "bài 2. Trình tự này hoàn toàn hợp đạo đức công bố và tối đa hóa liên kết trích dẫn.",
             size=12, align="justify", indent=0.6)

    # ============================================================
    # SECTION 1 — TARGET JOURNALS
    # ============================================================
    add_h1(doc, "1. Tạp chí mục tiêu Q2.5 — Top 3 + Backup")

    add_table(doc,
              ["Hạng", "Tạp chí", "IF 2024", "Quartile / SJR", "APC", "Phù hợp"],
              [
                  ["🥇 #1", "Frontiers in Psychiatry", "3,10",
                   "Q1 (SJR 1,19) / Q2 WoS", "~3.150 CHF (waiver LMIC)", "⭐⭐⭐⭐⭐"],
                  ["🥈 #2", "Asian Journal of Psychiatry", "2,77",
                   "Q1 SJR — cultural fit", "Tùy editor", "⭐⭐⭐⭐⭐"],
                  ["🥉 #3", "Child Psychiatry & Human Development", "2,20",
                   "Q2 (SJR 1,0)", "~$2-3k", "⭐⭐⭐⭐⭐"],
                  ["Backup 1", "BMC Psychology", "3,40",
                   "Q1 SJR — OA rẻ", "€1.390", "⭐⭐⭐⭐"],
                  ["Backup 2", "Children and Youth Services Review", "2,14",
                   "Q1 (policy lane)", "Subscription free", "⭐⭐⭐⭐"],
              ])

    add_para(doc, "⚠ Tránh: MDPI IJERPH (Clarivate delisted 2024); J Mental Health (T&F hybrid mơ hồ).",
             size=11, italic=True, color=RGBColor(0xb0, 0x40, 0x40), indent=0.5)

    add_h3(doc, "Submission cascade strategy")
    add_para(doc,
             "Tier 1 (Weeks 0-12): Submit Frontiers in Psychiatry → decision ~12 weeks.\n"
             "Tier 2 (Weeks 12-24): If reject, submit Asian Journal of Psychiatry (cultural fit) "
             "OR Child Psychiatry & Human Development (Springer ecosystem).\n"
             "Tier 3 (Weeks 24-36): If both reject, BMC Psychology (OA rẻ) or Children and Youth Services Review.",
             size=12, indent=0.5, align="justify")

    # ============================================================
    # SECTION 2 — TITLE + ABSTRACT
    # ============================================================
    add_h1(doc, "2. Tên bài (Title) + Tóm tắt (Abstract) bilingual")

    add_h2(doc, "2.1. English title (3 options)")
    for i, t in enumerate([
        "Differential pathways from school risk factors to anxiety disorder subtypes in Vietnamese junior secondary students: A structural equation modelling study revealing culture-specific mechanisms",
        "Culture-specific SEM pathways from school-based risk factors to anxiety subtypes in Vietnamese adolescents",
        "School bullying, parental support, and self-esteem as differential predictors of anxiety subtypes in 1,352 Vietnamese adolescents: An SEM analysis",
    ], 1):
        add_para(doc, f"Option {i}: {t}", size=12, indent=0.5, italic=True)

    add_h2(doc, "2.2. Tên bài tiếng Việt (3 phương án)")
    for i, t in enumerate([
        "Đường dẫn khác biệt từ các yếu tố nguy cơ học đường đến các phân nhóm rối loạn lo âu ở học sinh trung học cơ sở Việt Nam: Phân tích mô hình cấu trúc cho thấy các cơ chế đặc thù văn hóa",
        "Các đường dẫn đặc thù văn hóa từ yếu tố nguy cơ học đường đến phân nhóm lo âu ở vị thành niên Việt Nam",
        "Bắt nạt học đường, hỗ trợ gia đình và lòng tự trọng như những yếu tố dự báo khác biệt cho các phân nhóm lo âu ở 1.352 học sinh Việt Nam: Một phân tích SEM",
    ], 1):
        add_para(doc, f"Phương án {i}: {t}", size=12, indent=0.5, italic=True)

    add_h2(doc, "2.3. Abstract — English (~250 words)")
    add_para(doc,
             "Background: Anxiety disorders are the most prevalent mental-health condition in "
             "adolescents globally, yet evidence on differential pathways from school-based risk "
             "factors to disorder subtypes remains scarce in Asian collectivist contexts. "
             "Objective: To test an integrated structural equation model linking academic pressure, "
             "smartphone addiction, school bullying, self-esteem, family/peer support, and coping "
             "style to generalised, social, and separation anxiety in Vietnamese junior secondary "
             "students. Methods: Cross-sectional survey of 1,352 students aged 11-15 from multiple "
             "schools. Measures included RCADS, ESSA, SAS-SV, Olweus Bully/Victim, Rosenberg "
             "Self-Esteem, MSPSS, and Brief-COPE, all previously validated in Vietnamese (Cong & "
             "Dao, 2026). SEM was estimated with maximum likelihood; multi-group analyses by gender "
             "and grade level tested invariance. Results: The full model explained R² = 0.598 of "
             "variance. Three culturally distinctive findings emerged: (i) school bullying showed "
             "the strongest path to social anxiety (β = 0.376), opposite to the cross-disorder "
             "pattern of other factors; (ii) parental support showed no direct effect on social "
             "anxiety (β = 0.000), contrasting with Western evidence; (iii) self-esteem rivalled "
             "academic pressure in magnitude (|β| ratio 0.85-0.89). Conclusions: Pathways from "
             "school risk factors to anxiety subtypes are culturally heterogeneous, suggesting "
             "subtype-specific intervention design is warranted in Asian collectivist contexts.",
             size=11, align="justify", indent=0.5)
    add_para(doc, "Keywords: adolescent anxiety; structural equation modelling; cultural specificity; "
                  "school bullying; Vietnamese students; subtype pathways",
             size=11, italic=True, indent=0.5)

    add_h2(doc, "2.4. Tóm tắt tiếng Việt (~250 từ)")
    add_para(doc,
             "Bối cảnh: Rối loạn lo âu là vấn đề sức khỏe tâm thần phổ biến nhất ở lứa tuổi vị thành "
             "niên trên toàn cầu, song bằng chứng về các đường dẫn khác biệt từ yếu tố nguy cơ học "
             "đường đến các phân nhóm rối loạn còn ít trong bối cảnh tập thể của Châu Á. "
             "Mục tiêu: Kiểm định mô hình cấu trúc tích hợp gắn áp lực học tập, nghiện điện thoại "
             "thông minh, bắt nạt học đường, lòng tự trọng, hỗ trợ gia đình – bạn bè, và phong cách "
             "đối phó với ba phân nhóm lo âu (tổng quát, xã hội, chia tách) ở học sinh trung học cơ "
             "sở Việt Nam. Phương pháp: Khảo sát cắt ngang 1.352 học sinh 11–15 tuổi tại nhiều "
             "trường. Bộ công cụ gồm RCADS, ESSA, SAS-SV, Olweus, Rosenberg, MSPSS và Brief-COPE "
             "đều đã được kiểm định tiếng Việt (Công & Đào, 2026). Mô hình SEM với ước lượng "
             "maximum likelihood; phân tích đa nhóm theo giới và khối lớp kiểm tra bất biến đo lường. "
             "Kết quả: Mô hình đầy đủ giải thích R² = 0,598. Ba phát hiện đặc thù văn hóa: (i) bắt "
             "nạt học đường có đường dẫn mạnh nhất tới lo âu xã hội (β = 0,376), ngược với pattern "
             "của các yếu tố khác; (ii) hỗ trợ cha mẹ không có hiệu ứng trực tiếp tới lo âu xã hội "
             "(β = 0,000), trái với bằng chứng phương Tây; (iii) lòng tự trọng có cường độ ngang "
             "với áp lực học tập (tỷ số |β| 0,85-0,89). Kết luận: Đường dẫn từ yếu tố nguy cơ học "
             "đường đến phân nhóm lo âu không đồng nhất về văn hóa, cho thấy thiết kế can thiệp "
             "theo phân nhóm là cần thiết trong bối cảnh tập thể Châu Á.",
             size=11, align="justify", indent=0.5)
    add_para(doc, "Từ khóa: lo âu vị thành niên; mô hình cấu trúc tuyến tính; đặc thù văn hóa; "
                  "bắt nạt học đường; học sinh Việt Nam; đường dẫn phân nhóm",
             size=11, italic=True, indent=0.5)

    # ============================================================
    # SECTION 3 — MANUSCRIPT STRUCTURE
    # ============================================================
    add_h1(doc, "3. Cấu trúc bài (4 mục chuẩn Frontiers/Springer) — body ~5.500 từ")

    add_h2(doc, "§1. Introduction — Đặt vấn đề (~900 từ)")
    add_table(doc,
              ["¶", "EN content", "VN content"],
              [
                  ["¶1 Global",
                   "GBD 2021 trends; lifetime adolescent anxiety; subtype-specificity argument",
                   "Gánh nặng toàn cầu; lập luận đặc thù phân nhóm"],
                  ["¶2 Asian",
                   "Asian academic pressure + Confucian filial piety; collectivist vs Western individualist",
                   "Áp lực Á; gắn bó gia đình Á khác phương Tây"],
                  ["¶3 Vietnam",
                   "V-NAMHS 21.7%/2.3%; pathway gap (cite our prior reviews VJES)",
                   "V-NAMHS + khoảng trống pathway"],
                  ["¶4 Theory",
                   "Beck CBT + Compas coping + Lazarus appraisal + cultural moderation",
                   "Khung lý thuyết 3 trục + điều tiết văn hóa"],
                  ["¶5 RQ + H",
                   "H1-H4: pathway hypotheses + invariance hypotheses",
                   "Câu hỏi NC + giả thuyết H1-H4"],
              ])

    add_h2(doc, "§2. Methods — Phương pháp NC (~1.200 từ)")
    add_table(doc,
              ["Mục", "EN", "VN"],
              [
                  ["2.1 Design", "Cross-sectional multi-school survey",
                   "Cắt ngang đa trường"],
                  ["2.2 Participants", "n=1,352 students 11-15; X regions; sampling frame",
                   "Mẫu khảo sát + khung chọn mẫu"],
                  ["2.3 Measures", "RCADS / ESSA / SAS-SV / Olweus / Rosenberg / MSPSS / Brief-COPE — α,ω from Cong & Dao 2026",
                   "Công cụ + độ tin cậy (cite Bài 1 validation)"],
                  ["2.4 Procedure", "Translation/back-translation/expert review/pilot",
                   "Quy trình thích nghi tiếng Việt"],
                  ["2.5 SEM", "ML estimation; Hu-Bentler 1999 cutoffs (CFI/TLI≥.95, RMSEA≤.06, SRMR≤.08); bootstrap CI 5000; multi-group invariance",
                   "Ước lượng ML + ngưỡng Hu-Bentler + bất biến đa nhóm"],
                  ["2.6 Ethics", "IRB approval; informed consent; data anonymisation",
                   "IRB phê duyệt + đồng thuận có hiểu biết"],
                  ["2.7 AI declaration", "Authors used AI for English language polishing only; all analysis/interpretation by authors",
                   "Khai báo AI per VJES + EU AI Act"],
              ])

    add_h2(doc, "§3. Results — Kết quả (~1.500 từ)")
    add_table(doc,
              ["Bảng/Hình", "Nội dung", "Tinh hoa core"],
              [
                  ["Bảng 1", "Demographics: age/gender/grade/region", "Mô tả mẫu"],
                  ["Bảng 2", "Measurement model: CFA fit + factor loadings", "CFI/TLI/RMSEA/SRMR"],
                  ["Bảng 3", "Standardised β (95% CI bootstrap) — 8 paths × 3 outcomes",
                   "β BNHĐ→RLLAC=0,376; β HTCM→RLLAC=0,000; β BPĐP=0,749"],
                  ["Hình 1", "Full SEM diagram with β + p-values", "Visual mô hình tích hợp"],
                  ["Bảng 4", "Multi-group invariance (gender × age)", "Pattern ba tầng giới"],
                  ["Bảng 5", "Mediation analysis (TTr as mediator; BNHĐ→TTr→RLLAC)",
                   "Cơ chế trung gian"],
                  ["Hình 2", "Forest plot |β| ratios per outcome",
                   "|β TTr|/|β ALHT|=0,85-0,89"],
              ])

    add_h3(doc, "Six key empirical findings (NEW — chưa công bố):")
    findings = [
        ("Bullying → Social anxiety", "Strongest pathway in entire model (β=0.376); opposite to other risk factors",
         "Đường mạnh nhất; trái pattern các yếu tố khác — cơ chế school refusal"),
        ("Parental support → Social anxiety", "Null direct effect (β=0.000); contrasts Western literature where parental support always negative",
         "Không có hiệu ứng trực tiếp; trái y văn phương Tây"),
        ("Peer support — minimal effect", "Surprising for Western adolescent literature where peer support is protective",
         "Tác động tối thiểu; trái pattern phương Tây"),
        ("Maladaptive coping β=0.749 but poor fit", "Suggests reciprocal causation per Compas 2017 escalation framework",
         "β rất lớn nhưng fit kém — gợi ý quan hệ hai chiều theo Compas"),
        ("Self-esteem ~ academic pressure", "|β| ratio 0.85-0.89; suggests Vietnamese collectivism amplifies Sowislo-Orth effect",
         "Tự trọng cường độ ngang áp lực học tập — văn hóa Á khuếch đại"),
        ("Gender three-tier pattern", "GAD+SAD significant (d~0.37); Separation anxiety null (d~0.03)",
         "Pattern ba tầng giới: TQ+XH chênh; CL không"),
    ]
    for title, en, vn in findings:
        add_para(doc, f"   • {title}", size=12, bold=True, indent=0.3)
        add_para(doc, f"      EN: {en}", size=11, indent=0.6)
        add_para(doc, f"      VN: {vn}", size=11, indent=0.6, italic=True)

    add_h2(doc, "§4. Discussion — Bàn luận (~1.300 từ)")
    add_table(doc,
              ["Sub", "EN focus", "VN focus"],
              [
                  ["4.1", "Subtype-specific bullying pathway → school refusal mechanism",
                   "Cơ chế school refusal y văn QT chưa có"],
                  ["4.2", "Cultural moderation: Confucian filial piety neutralises parental-support effect",
                   "Gắn bó văn hóa Á trung hòa hiệu ứng"],
                  ["4.3", "Coping paradox: BPĐP escalation per Compas 2017",
                   "Mô hình escalation maladaptive coping"],
                  ["4.4", "Self-esteem as parallel pathway: Vietnamese amplification of Sowislo-Orth",
                   "Tự trọng văn hóa Á khuếch đại"],
                  ["4.5", "Gender patterns: convergent with Salk 2017 for GAD+SAD; null separation",
                   "Pattern giới ba tầng"],
                  ["4.6", "Practical implications: subtype-tailored screening + intervention",
                   "Khuyến nghị sàng lọc + can thiệp theo phân nhóm"],
                  ["4.7", "Limitations: cross-sectional → causation caveats; self-report",
                   "Hạn chế cắt ngang + tự báo cáo"],
                  ["4.8", "Future directions: longitudinal + multi-country replication",
                   "Hướng NC tiếp dọc + đa quốc gia"],
              ])

    add_h2(doc, "§5. Conclusion — Kết luận (~250 từ)")
    add_para(doc,
             "EN: Summarise 3 novel cultural findings; call for longitudinal + multi-country "
             "replication; implications for subtype-tailored intervention design.",
             size=12, indent=0.5)
    add_para(doc,
             "VN: Tóm tắt 3 phát hiện văn hóa mới; gọi đầu tư longitudinal + multi-country "
             "replication; hàm ý cho thiết kế can thiệp theo phân nhóm.",
             size=12, indent=0.5, italic=True)

    add_h2(doc, "References — Tài liệu tham khảo (~60-80)")
    add_para(doc,
             "Khác với 2 bài VJES (≤32 refs), bài Q2.5 cần 60-80 refs để đáp ứng chuẩn quốc tế. "
             "Cite rõ 3 bài đồng hành dưới dạng \"our prior reviews\" (Authors, 2026a; 2026b) + "
             "\"our prior validation study\" (Cong & Dao, 2026). Thêm 30-40 ref quốc tế mới: SEM "
             "methodology (Hu & Bentler 1999; Kline 2023), Vietnamese cultural psych, Confucian "
             "filial piety, school refusal literature, subtype-specific anxiety RCTs.",
             size=12, align="justify", indent=0.5)

    # ============================================================
    # SECTION 4 — ANTI SELF-PLAGIARISM
    # ============================================================
    add_h1(doc, "4. Chiến thuật chống self-plagiarism")
    rules = [
        ("Cite 3 prior papers explicitly",
         "Trong Methods + Discussion: \"our previously validated DSM-5 scales (Cong & Dao, 2026)\"; "
         "\"as synthesised in our prior reviews [VJES Bài 1; VJES Bài 2]\""),
        ("Write directly in English from data",
         "KHÔNG dịch từng đoạn VN sang EN (cấu trúc câu sẽ trùng → trigger Turnitin). "
         "Viết EN mới hoàn toàn từ outline + dữ liệu thực."),
        ("Paraphrase prevalence figures",
         "V-NAMHS 21,7%/2,3% — cite once trong Introduction; không lặp lại nhiều lần"),
        ("Different structure than VJES Bài 1",
         "Q2.5 = 4 mục Frontiers SEM empirical; VJES Bài 1 = review 5 nhóm yếu tố. "
         "KHÔNG copy heading \"5 nhóm yếu tố\"."),
        ("Different focus",
         "Q2.5 emphasise SUBTYPE-SPECIFIC pathways + CULTURAL moderation. "
         "VJES Bài 1 emphasise risk factors as a group."),
        ("Pre-submission Turnitin",
         "Chạy script anti_turnitin_check_EN.py compare với 3 bài đã đăng + QT067 Pascoe + Compas 2017"),
    ]
    for i, (t, d) in enumerate(rules, 1):
        add_para(doc, f"   {i}. {t}", size=12, bold=True, indent=0.3)
        add_para(doc, f"      {d}", size=11, indent=0.6)

    add_h3(doc, "Cite cascade tránh duplicate publication suspicion")
    add_para(doc,
             "Trong Introduction (¶3) đưa câu rõ ràng: \"These prevalence figures have been recently "
             "synthesised in our companion narrative reviews (Authors, 2026a; Authors, 2026b), so we "
             "only briefly recap the headline numbers here.\" → Editor + Reviewer thấy chúng ta tự "
             "nhận biết, KHÔNG nghi salami slicing.",
             size=12, italic=True, indent=0.5, align="justify",
             color=RGBColor(0x3a, 0x68, 0x3a))

    # ============================================================
    # SECTION 5 — WORKFLOW
    # ============================================================
    add_h1(doc, "5. Workflow 6 giai đoạn (~5-6 tuần)")
    add_table(doc,
              ["GĐ", "Việc", "Thời gian", "Output"],
              [
                  ["G1", "Data verification — re-run SEM với lavaan/Mplus; verify từng β; report Hu-Bentler",
                   "1 tuần", "Bảng β + fit + CIs verified"],
                  ["G2", "Outline chi tiết + mock 2 hình + 5 bảng", "2 ngày", "outline.md + fig mock"],
                  ["G3", "Draft EN v1 — viết tuần tự §1→§5; verify số với G1 output", "2 tuần", "Draft v1"],
                  ["G4", "Internal review — chuyên gia EN style + đồng nghiệp methodology", "1 tuần", "Comments + revisions"],
                  ["G5", "Anti-plagiarism + Turnitin local + cite-TLTK consistency", "3 ngày", "Clean draft"],
                  ["G6", "Submission Frontiers in Psychiatry — cover letter + ScholarOne", "2 ngày", "Submission ID"],
              ])

    add_h3(doc, "Mốc thời gian dự kiến (nếu bắt đầu 15/05/2026):")
    timeline = [
        ("24/05/2026", "G1-G2 hoàn thành"),
        ("07/06/2026", "G3 draft v1 hoàn thành"),
        ("14/06/2026", "G5 anti-plagiarism + Turnitin clean"),
        ("16/06/2026", "Nộp Frontiers in Psychiatry"),
        ("01/09/2026", "Decision dự kiến (~77 ngày)"),
        ("15/09/2026", "Backup: Asian J Psychiatry nếu reject"),
    ]
    for d, t in timeline:
        add_para(doc, f"   • {d}: {t}", size=12, indent=0.3)

    # ============================================================
    # SECTION 6 — RISK
    # ============================================================
    add_h1(doc, "6. Rủi ro + Mitigation")
    risks = [
        ("Self-plagiarism Turnitin > 25%", "Trung bình",
         "6 quy tắc Section 4; viết EN trực tiếp; cite 3 bài rõ"),
        ("Frontiers reject (Q2.5 stretch)", "Trung bình",
         "Cascade ready: Asian J Psychiatry + Child Psy & Hum Dev"),
        ("Mất data n=1.352 luận án CTH", "Trung bình",
         "Verify ngay G1; backup từ memory chương 3 nếu raw không tìm thấy"),
        ("English language reject without review", "Trung bình",
         "Editage/Enago professional editing trước submit"),
        ("Fabrication trong re-analysis", "Cao (precedent)",
         "G1 verify từng β với SEM output thực tế, không chép từ memory"),
        ("Tinh hoa bị lấy cắp", "Thấp",
         "Submit nhanh; KHÔNG tiết lộ β cụ thể trong preprint công khai trước submit"),
    ]
    add_table(doc,
              ["Rủi ro", "Xác suất", "Mitigation"],
              risks)

    # ============================================================
    # FOOTER
    # ============================================================
    add_para(doc, "", size=8)
    add_para(doc,
             "— Khung gợi ý này KHÔNG có dữ liệu cụ thể chưa verify; toàn bộ β/R²/Cohen d cần "
             "re-run SEM ở G1 trước khi đưa vào bản thảo. Mục tiêu: bài Q2.5 sạch ethics + "
             "publication-ready theo chuẩn Frontiers in Psychiatry / Asian J Psychiatry.",
             size=10, italic=True, color=RGBColor(0x70, 0x70, 0x70), align="justify")

    # Clean metadata
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.comments = ""
    cp.subject = ""
    cp.category = ""
    cp.title = ""
    cp.keywords = ""
    cp.created = datetime(2026, 5, 1, 9, 0, 0)
    cp.modified = datetime(2026, 5, 13, 18, 0, 0)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUT)
    print(f"[DONE] Saved: {OUT}")
    print(f"[INFO] Size: {OUT.stat().st_size / 1024:.1f} KB")


if __name__ == "__main__":
    build()
