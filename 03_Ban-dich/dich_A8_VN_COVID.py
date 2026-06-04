# -*- coding: utf-8 -*-
"""Dịch đầy đủ A8 — Nguyen et al. 2023 — Medicine Q1 — Lo âu VN COVID n=5.730"""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.1097/MD.0000000000033559', size=10)
add_p(doc, 'PMC: https://pmc.ncbi.nlm.nih.gov/articles/PMC10118042/', size=10)

add_heading(doc, 'Lo \u00e2u v\u00e0 c\u00e1c y\u1ebfu t\u1ed1 li\u00ean quan \u1edf sinh vi\u00ean Vi\u1ec7t Nam trong \u0111\u1ea1i d\u1ecbch COVID-19: Nghi\u00ean c\u1ee9u c\u1eaft ngang', 1)
h = doc.add_paragraph()
r = h.add_run('Anxiety and associated factors among Vietnamese students during COVID-19 pandemic: A cross-sectional study')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'TH\u00d4NG TIN TH\u01af M\u1ee4C', 2)
add_info_table(doc, [
    ('Ti\u00eau \u0111\u1ec1 g\u1ed1c', 'Anxiety and associated factors among Vietnamese students during COVID-19 pandemic: A cross-sectional study'),
    ('Ti\u00eau \u0111\u1ec1 d\u1ecbch', 'Lo \u00e2u v\u00e0 c\u00e1c y\u1ebfu t\u1ed1 li\u00ean quan \u1edf sinh vi\u00ean Vi\u1ec7t Nam trong \u0111\u1ea1i d\u1ecbch COVID-19: Nghi\u00ean c\u1ee9u c\u1eaft ngang'),
    ('T\u00e1c gi\u1ea3', 'Long Xuan Nguyen, Linh Thi Dieu Dao*, Anh Nhat Ta, Ha Thanh Le, Hieu Van Nguyen, Lan Thi Mai Nguyen, Phuong Thi Nguyen, Thang Thi Nguyen, Thanh Chi Ta, Thuong Hiep Nguyen, Trang Thi Nguyen, Tuan Anh Huynh, Anh Quynh Hoang, Linh Thi Hanh Duong, Ly Hoang Do, Nam Tien Pham (16 t\u00e1c gi\u1ea3)'),
    ('C\u01a1 quan', '\u0110H Ngo\u1ea1i ng\u1eef \u2014 \u0110H Qu\u1ed1c gia H\u00e0 N\u1ed9i (ch\u00ednh); \u0110H S\u01b0 ph\u1ea1m H\u00e0 N\u1ed9i 2; \u0110H S\u01b0 ph\u1ea1m H\u00e0 N\u1ed9i; Vi\u1ec7n T\u00e2m l\u00fd h\u1ecdc VASS; \u0110H Th\u1ee7 \u0111\u00f4; T\u1ed5ng Li\u00ean \u0111o\u00e0n Lao \u0111\u1ed9ng VN; FPT Polytechnic BMT; Deakin University \u00dac; \u0110H YTCC H\u00e0 N\u1ed9i'),
    ('T\u1ea1p ch\u00ed', 'Medicine (Baltimore) \u2014 Wolters Kluwer / LWW (Q1, IF \u2248 1,6)'),
    ('Xu\u1ea5t b\u1ea3n', '2023, Vol. 102(16), e33559, 6 trang'),
    ('DOI', '10.1097/MD.0000000000033559'),
    ('Lo\u1ea1i NC', 'C\u1eaft ngang \u2014 kh\u1ea3o s\u00e1t tr\u1ef1c tuy\u1ebfn'),
    ('M\u1eabu', '5.730 sinh vi\u00ean Vi\u1ec7t Nam to\u00e0n qu\u1ed1c, 8\u20139/2021 (l\u00e0n s\u00f3ng 4 COVID-19)'),
    ('C\u00f4ng c\u1ee5', 'GAD-7 (Generalized Anxiety Disorder Questionnaire, \u03b1 = 0,92); ng\u01b0\u1ee1ng \u226510'),
    ('\u0110\u1ea1o \u0111\u1ee9c', 'H\u1ed9i T\u00e2m l\u00fd h\u1ecdc Vi\u1ec7t Nam, Q\u0110 s\u1ed1 10/2021/HTLHVN-DD'),
    ('Truy c\u1eadp', 'Open Access \u2014 CC BY-NC 4.0'),
])
add_page_ref(doc, '1\u20136', 'Medicine', 'Vol. 102(16), 2023')

# TOM TAT
add_heading(doc, 'T\u00d3M T\u1eaeT', 2)
add_p(doc, 'Ch\u0103m s\u00f3c s\u1ee9c kh\u1ecfe t\u00e2m th\u1ea7n cho sinh vi\u00ean n\u00f3i chung, \u0111\u1eb7c bi\u1ec7t l\u00e0 lo \u00e2u, l\u00e0 v\u1ea5n \u0111\u1ec1 \u0111\u00e1ng k\u1ec3 c\u1ea7n \u0111\u01b0\u1ee3c quan t\u00e2m h\u01a1n, \u0111\u1eb7c bi\u1ec7t trong \u0111\u1ea1i d\u1ecbch COVID-19. Nghi\u00ean c\u1ee9u n\u00e0y nh\u1eb1m m\u00f4 t\u1ea3 t\u1ef7 l\u1ec7 lo \u00e2u v\u00e0 ki\u1ec3m tra c\u00e1c y\u1ebfu t\u1ed1 li\u00ean quan \u1edf sinh vi\u00ean trong \u0111\u1ea1i d\u1ecbch COVID-19 t\u1ea1i Vi\u1ec7t Nam.')

p = doc.add_paragraph()
r = p.add_run('Ph\u01b0\u01a1ng ph\u00e1p: Nghi\u00ean c\u1ee9u c\u1eaft ngang tr\u1ef1c tuy\u1ebfn t\u1eeb th\u00e1ng 8 \u0111\u1ebfn th\u00e1ng 9/2021 tr\u00ean 5.730 sinh vi\u00ean. Kh\u1ea3o s\u00e1t tr\u1ef1c tuy\u1ebfn thu th\u1eadp th\u00f4ng tin nh\u00e2n kh\u1ea9u x\u00e3 h\u1ed9i, v\u00e0 B\u1ea3ng h\u1ecfi R\u1ed1i lo\u1ea1n Lo \u00e2u T\u1ed5ng qu\u00e1t 7 m\u1ee5c (GAD-7, Cronbach \u03b1 = 0,92) \u0111\u01b0\u1ee3c s\u1eed d\u1ee5ng \u0111\u1ec3 \u0111\u00e1nh gi\u00e1 tri\u1ec7u ch\u1ee9ng lo \u00e2u. Ng\u01b0\u1ee1ng c\u1eaft: GAD-7 \u2265 10 (v\u1eeba\u2013n\u1eb7ng). H\u1ed3i quy Poisson m\u1ea1nh m\u1ebd (Robust Poisson regression) t\u00ednh T\u1ef7 s\u1ed1 T\u1ef7 l\u1ec7 (PR) v\u00e0 KTC 95%.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('K\u1ebft qu\u1ea3: T\u1ef7 l\u1ec7 lo \u00e2u: 16,2% (KTC 95%: 15,3\u201317,2%). N\u1eef cao h\u01a1n nam: 17,9% vs 13,1% (p < 0,001). Y\u1ebfu t\u1ed1 li\u00ean quan (Robust Poisson PR): Nam th\u1ea5p h\u01a1n n\u1eef (PR = 0,66; KTC: 0,58\u20130,74); K\u00fd t\u00fac x\u00e1 (PR = 1,71; KTC: 1,39\u20132,12); COVID nhi\u1ec5m (PR = 2,29; KTC: 1,40\u20133,75); Ti\u00eam v\u1eafc-xin b\u1ea3o v\u1ec7 (li\u1ec1u 1: PR = 0,97; li\u1ec1u 2: PR = 0,95); S\u1ee9c kh\u1ecfe k\u00e9m (PR = 1,71; KTC: 1,34\u20132,18); H\u1ecdc t\u1eadp k\u00e9m (PR = 1,44; KTC: 1,12\u20131,87); Quan h\u1ec7 XH k\u00e9m (PR = 1,54; KTC: 1,16\u20132,04).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

add_p(doc, 'K\u1ebft lu\u1eadn: M\u1ed9t s\u1ed1 l\u01b0\u1ee3ng \u0111\u00e1ng k\u1ec3 sinh vi\u00ean tr\u1ea3i qua lo \u00e2u trong COVID-19, li\u00ean quan \u0111\u1ebfn nhi\u1ec1u y\u1ebfu t\u1ed1. C\u1ea7n can thi\u1ec7p t\u00e2m l\u00fd h\u1ed7 tr\u1ee3 sinh vi\u00ean trong v\u00e0 sau \u0111\u1ea1i d\u1ecbch v\u00e0 c\u00e1c kh\u1ee7ng ho\u1ea3ng s\u1ee9c kh\u1ecfe kh\u00e1c.')

# DANH GIA NHANH
add_heading(doc, 'T\u00d3M T\u1eaeT \u0110\u00c1NH GI\u00c1 NHANH', 2)
for b in [
    'Medicine Q1 (Wolters Kluwer/LWW) \u2014 t\u1ea1p ch\u00ed QU\u1ed0C T\u1ebe c\u00f3 ch\u1ec9 m\u1ee5c PubMed.',
    'M\u1eabu L\u1edaN NH\u1ea4T trong c\u00e1c NC lo \u00e2u VN d\u00f9ng GAD-7: n = 5.730 (so: Hoa 2024 VN01 n = 3.910).',
    'GAD-7 \u2014 c\u00f4ng c\u1ee5 chu\u1ea9n qu\u1ed1c t\u1ebf, \u03b1 = 0,92 (r\u1ea5t cao). Ng\u01b0\u1ee1ng \u226510 \u2014 ch\u00ednh x\u00e1c: sensitivity + specificity > 80%.',
    'Lo \u00e2u 16,2% \u2014 TH\u1ea4P H\u01a0N Hoa 2024 (40,6%, GAD-7 \u22655). Do ng\u01b0\u1ee1ng cao h\u01a1n (\u226510 vs \u22655).',
    'Robust Poisson regression \u2014 ph\u01b0\u01a1ng ph\u00e1p ph\u00f9 h\u1ee3p cho d\u1eef li\u1ec7u t\u1ef7 l\u1ec7 (PR thay v\u00ec OR). \u00cdt NC VN d\u00f9ng.',
    'B\u1ed1i c\u1ea3nh COVID-19 l\u00e0n s\u00f3ng 4 (8\u20139/2021) \u2014 giai \u0111o\u1ea1n n\u1eb7ng nh\u1ea5t t\u1ea1i VN.',
    '16 t\u00e1c gi\u1ea3 t\u1eeb 9 c\u01a1 quan \u2014 h\u1ee3p t\u00e1c \u0111a ngành.',
    'Open Access CC BY-NC 4.0.',
]:
    add_p(doc, '\u2022 ' + b)
add_p(doc, 'H\u1ea1n ch\u1ebf:', bold=True)
for b in [
    'Ch\u1ec9 SINH VI\u00caN (\u226518 tu\u1ed5i) \u2014 kh\u00f4ng bao g\u1ed3m HS THCS/THPT. So: \u0111\u1ec1 t\u00e0i t\u1eadp trung c\u1ea5p 2+3.',
    'L\u1ea5y m\u1eabu t\u1ef1 ch\u1ecdn (self-selection) tr\u1ef1c tuy\u1ebfn \u2014 thi\u00ean l\u1ec7ch ch\u1ecdn.',
    'C\u1eaft ngang \u2014 kh\u00f4ng nh\u00e2n qu\u1ea3.',
    'B\u1ed1i c\u1ea3nh COVID \u0111\u1eb7c bi\u1ec7t \u2014 k\u1ebft qu\u1ea3 c\u00f3 th\u1ec3 kh\u00f4ng kh\u00e1i qu\u00e1t cho giai \u0111o\u1ea1n b\u00ecnh th\u01b0\u1eddng.',
    'GAD-7 \u226510 (v\u1eeba\u2013n\u1eb7ng) \u2014 b\u1ecf s\u00f3t lo \u00e2u nh\u1eb9 (5\u20139). Hoa 2024 d\u00f9ng \u22655: 40,6%.',
]:
    add_p(doc, '\u2022 ' + b)

# 1. GIOI THIEU
add_page_ref(doc, '1\u20132', 'Medicine', 'Vol. 102(16), 2023')
add_heading(doc, '1. GI\u1edaI THI\u1ec6U', 2)
add_p(doc, 'Tr\u01b0\u1edbc khi \u0111\u1ea1i d\u1ecbch COVID-19 x\u1ea3y ra, c\u00e1c v\u1ea5n \u0111\u1ec1 SKTT c\u1ee7a sinh vi\u00ean \u0111\u00e3 \u0111\u01b0\u1ee3c x\u00e1c \u0111\u1ecbnh l\u00e0 tr\u1edf ng\u1ea1i h\u00e0ng \u0111\u1ea7u \u0111\u1ed1i v\u1edbi th\u00e0nh t\u00edch h\u1ecdc t\u1eadp. R\u1ed1i lo\u1ea1n t\u00e2m th\u1ea7n c\u00f3 th\u1ec3 \u1ea3nh h\u01b0\u1edfng \u0111\u1ebfn \u0111\u1ed9ng l\u1ef1c, t\u1eadp trung v\u00e0 t\u01b0\u01a1ng t\u00e1c x\u00e3 h\u1ed9i. Theo b\u00e1o c\u00e1o c\u1ee7a Collegiate Mental Health (2019) bao g\u1ed3m 163 trung t\u00e2m t\u01b0 v\u1ea5n t\u1eeb 54 qu\u1ed1c gia, lo \u00e2u l\u00e0 v\u1ea5n \u0111\u1ec1 SKTT ph\u1ed5 bi\u1ebfn nh\u1ea5t, v\u1edbi 82.685 sinh vi\u00ean (62,7% m\u1eabu) b\u00e1o c\u00e1o c\u00f3 tri\u1ec7u ch\u1ee9ng lo \u00e2u.')
add_p(doc, 'Hai n\u0103m sau khi COVID-19 xu\u1ea5t hi\u1ec7n, WHO g\u1ecdi COVID-19 l\u00e0 "\u0111\u1ea1i d\u1ecbch g\u00e2y lo ng\u1ea1i r\u1ed9ng r\u00e3i v\u00e0 l\u00e0m t\u0103ng lo \u00e2u v\u00e0 c\u0103ng th\u1eb3ng to\u00e0n c\u1ea7u." T\u1ed5ng quan h\u1ec7 th\u1ed1ng 36 NC (9/2020\u20132/2021) cho th\u1ea5y t\u1ef7 l\u1ec7 lo \u00e2u kho\u1ea3ng 41% (KTC 95% = 0,34\u20130,49).')
add_p(doc, 'Jiang et al. (2021) kh\u1ea3o s\u00e1t 1.195 SV t\u1eeb 4 n\u01b0\u1edbc ch\u00e2u \u00c1\u2013Th\u00e1i B\u00ecnh D\u01b0\u01a1ng (Malaysia, Indonesia, Th\u00e1i Lan, Trung Qu\u1ed1c): lo \u00e2u l\u00e0 v\u1ea5n \u0111\u1ec1 t\u00e2m l\u00fd ph\u1ed5 bi\u1ebfn nh\u1ea5t, ti\u1ebfp theo l\u00e0 tr\u1ea7m c\u1ea3m v\u00e0 stress. Lo \u00e2u n\u1eb7ng v\u00e0 c\u1ef1c k\u1ef3 n\u1eb7ng \u0111\u01b0\u1ee3c quan s\u00e1t \u1edf 20,5% ng\u01b0\u1eddi tham gia.')
add_p(doc, 'T\u1ea1i M\u1ef9, kh\u1ea3o s\u00e1t 2.031 SV cho th\u1ea5y 38,48% c\u00f3 lo \u00e2u m\u1ee9c v\u1eeba\u2013n\u1eb7ng, 71,26% t\u0103ng c\u0103ng th\u1eb3ng v\u00e0 lo \u00e2u trong \u0111\u1ea1i d\u1ecbch. T\u1ea1i Bangladesh, 14,8% tr\u1ea3i qua lo \u00e2u c\u1ef1c \u0111\u1ed9, 82,5% nh\u1eb9\u2013c\u1ef1c \u0111\u1ed9.')
add_p(doc, 'T\u1ea1i Vi\u1ec7t Nam, m\u1ed9t s\u1ed1 \u00edt NC \u0111\u00e3 \u0111\u01b0\u1ee3c th\u1ef1c hi\u1ec7n v\u1ec1 SKTT trong COVID-19 \u1edf d\u00e2n s\u1ed1 chung, nh\u00e2n vi\u00ean y t\u1ebf, v\u00e0 sinh vi\u00ean. Hi\u1ec3u bi\u1ebft t\u00ecnh h\u00ecnh lo \u00e2u SV v\u00e0 c\u00e1c y\u1ebfu t\u1ed1 li\u00ean quan d\u01b0\u1edbi t\u00e1c \u0111\u1ed9ng COVID-19 l\u00e0 v\u1ea5n \u0111\u1ec1 c\u1ea5p thi\u1ebft. NC n\u00e0y nh\u1eb1m m\u00f4 t\u1ea3 t\u1ef7 l\u1ec7 lo \u00e2u v\u00e0 ki\u1ec3m tra y\u1ebfu t\u1ed1 li\u00ean quan \u1edf SV Vi\u1ec7t Nam trong COVID-19.')

# 2. PHUONG PHAP
add_page_ref(doc, '2\u20133', 'Medicine', 'Vol. 102(16), 2023')
add_heading(doc, '2. PH\u01af\u01a0NG PH\u00c1P', 2)
add_p(doc, '2.1. Thi\u1ebft k\u1ebf v\u00e0 ng\u01b0\u1eddi tham gia', bold=True)
add_p(doc, 'Kh\u1ea3o s\u00e1t c\u1eaft ngang tr\u1ef1c tuy\u1ebfn. Ti\u00eau ch\u00ed: \u226518 tu\u1ed5i, s\u1ed1ng t\u1ea1i VN, \u0111\u1ed3ng \u00fd tham gia, truy c\u1eadp b\u1ea3ng h\u1ecfi tr\u1ef1c tuy\u1ebfn.')
add_p(doc, '2.2. M\u1eabu v\u00e0 l\u1ea5y m\u1eabu', bold=True)
add_p(doc, 'Kh\u1ea3o s\u00e1t trong 2 tu\u1ea7n gi\u00e3n c\u00e1ch x\u00e3 h\u1ed9i (23/8\u20135/9/2021). L\u1ea5y m\u1eabu t\u1ef1 ch\u1ecdn (self-selection): link b\u1ea3ng h\u1ecfi g\u1eedi qua m\u1ea1ng l\u01b0\u1edbi gi\u1ea3ng vi\u00ean, chuy\u1ec3n ti\u1ebfp cho SV qua Email, Facebook, Zalo. SV c\u0169ng \u0111\u01b0\u1ee3c khuy\u1ebfn kh\u00edch chia s\u1ebb r\u1ed9ng r\u00e3i. T\u1ed5ng c\u1ed9ng 5.730 SV tham gia.')
add_p(doc, '2.3. C\u00f4ng c\u1ee5 \u0111o', bold=True)
add_p(doc, 'Ph\u1ea7n 1: Th\u00f4ng tin nh\u00e2n kh\u1ea9u. Ph\u1ea7n 2: GAD-7 (Spitzer et al., 2006) \u2014 7 c\u00e2u, Likert 4 \u0111i\u1ec3m (0\u20133), \u0111\u00e1nh gi\u00e1 lo \u00e2u 2 tu\u1ea7n qua. T\u1ed5ng 0\u201321. Ph\u00e2n nh\u00f3m: kh\u00f4ng (0\u20134), nh\u1eb9 (5\u20139), v\u1eeba (10\u201314), n\u1eb7ng (15\u201321). Ng\u01b0\u1ee1ng d\u01b0\u01a1ng t\u00ednh: \u226510 (sensitivity + specificity > 80%). Phi\u00ean b\u1ea3n ti\u1ebfng Vi\u1ec7t \u0111\u00e3 \u0111\u01b0\u1ee3c s\u1eed d\u1ee5ng trong nhi\u1ec1u NC t\u1ea1i VN. Cronbach \u03b1 = 0,92 trong NC n\u00e0y.')
add_p(doc, '2.4. Ph\u00e2n t\u00edch', bold=True)
add_p(doc, 'Chi-square so s\u00e1nh \u0111\u1eb7c \u0111i\u1ec3m. H\u1ed3i quy Poisson \u0111a bi\u1ebfn v\u1edbi sai s\u1ed1 m\u1ea1nh m\u1ebd (Robust Poisson regression) \u2014 t\u00ednh T\u1ef7 s\u1ed1 T\u1ef7 l\u1ec7 (PR) v\u00e0 KTC 95%. M\u1ee9c \u00fd ngh\u0129a p < 0,05. Stata 14.2.')

# 3. KET QUA
add_page_ref(doc, '3\u20134', 'Medicine', 'Vol. 102(16), 2023')
add_heading(doc, '3. K\u1ebeT QU\u1ea2', 2)
add_p(doc, '3.1. \u0110\u1eb7c \u0111i\u1ec3m m\u1eabu (n = 5.730)', bold=True)
add_p(doc, 'N\u1eef 63,9% (3.667). S\u1ed1ng c\u00f9ng gia \u0111\u00ecnh 81,8% (4.687). Ch\u01b0a ti\u1ebfp x\u00fac COVID 88,5% (5.073). Ch\u01b0a ti\u00eam v\u1eafc-xin 75,5% (4.328). S\u1ee9c kh\u1ecfe \u1ed5n \u0111\u1ecbnh 84,8% (4.857). H\u1ecdc t\u1eadp \u1ed5n \u0111\u1ecbnh 50,0% (2.866). Quan h\u1ec7 XH \u1ed5n \u0111\u1ecbnh 55,4% (3.174).')

add_p(doc, '3.2. T\u1ef7 l\u1ec7 lo \u00e2u', bold=True)

add_heading(doc, 'B\u1ea3ng 1. T\u1ef7 l\u1ec7 lo \u00e2u GAD-7 \u226510 theo nh\u00f3m', 3)
add_table(doc,
    ['Nh\u00f3m', 'Lo \u00e2u (%)', 'Kh\u00f4ng lo \u00e2u (%)', 'p'],
    [['T\u1ed5ng', '16,2% (929)', '83,8% (4.801)', ''],
     ['N\u1eef', '17,9% (658)', '82,1% (3.009)', '<0,001'],
     ['Nam', '13,1% (271)', '86,9% (1.792)', ''],
     ['S\u1ed1ng c\u00f9ng gia \u0111\u00ecnh', '14,9% (697)', '85,1% (3.990)', '<0,001'],
     ['K\u00fd t\u00fac x\u00e1', '32,4% (59)', '67,6% (123)', ''],
     ['\u1ede m\u1ed9t m\u00ecnh tr\u1ecd', '24,8% (75)', '75,2% (227)', ''],
     ['COVID nhi\u1ec5m', '37,5% (9)', '62,5% (15)', '<0,001'],
     ['COVID ti\u1ebfp x\u00fac F0', '30,9% (17)', '69,1% (38)', ''],
     ['Ch\u01b0a ti\u1ebfp x\u00fac', '15,1% (767)', '84,9% (4.306)', ''],
     ['S\u1ee9c kh\u1ecfe k\u00e9m h\u01a1n', '25,1% (107)', '74,9% (320)', '<0,001'],
     ['H\u1ecdc t\u1eadp k\u00e9m h\u01a1n', '20,1% (486)', '79,9% (1.928)', '<0,001'],
     ['Quan h\u1ec7 XH k\u00e9m h\u01a1n', '20,5% (445)', '79,5% (1.724)', '<0,001']],
    widths=[4.0, 3.0, 3.0, 1.5])
add_p(doc, 'K\u00fd t\u00fac x\u00e1 = 32,4% lo \u00e2u \u2014 g\u1ea5p \u0111\u00f4i s\u1ed1ng c\u00f9ng gia \u0111\u00ecnh (14,9%). COVID nhi\u1ec5m = 37,5% \u2014 cao nh\u1ea5t.', size=9, italic=True)

add_p(doc, '3.3. Y\u1ebfu t\u1ed1 li\u00ean quan (Robust Poisson regression)', bold=True)

add_heading(doc, 'B\u1ea3ng 2. Y\u1ebfu t\u1ed1 li\u00ean quan lo \u00e2u \u2014 Robust Poisson PR (n = 5.730)', 3)
add_table(doc,
    ['Y\u1ebfu t\u1ed1', 'PR', 'KTC 95%', '\u00dd ngh\u0129a'],
    [['Gi\u1edbi t\u00ednh nam (vs n\u1eef)', '0,66', '0,58\u20130,74', 'Nam TH\u1ea4P h\u01a1n n\u1eef 34%'],
     ['K\u00fd t\u00fac x\u00e1 (vs gia \u0111\u00ecnh)', '1,71', '1,39\u20132,12', 'T\u0102NG 71%'],
     ['\u1ede m\u1ed9t m\u00ecnh (vs gia \u0111\u00ecnh)', '1,36', '1,10\u20131,69', 'T\u0102NG 36%'],
     ['COVID nhi\u1ec5m (vs kh\u00f4ng)', '2,29', '1,40\u20133,75', 'T\u0102NG 129%'],
     ['Ti\u1ebfp x\u00fac F0 (vs kh\u00f4ng)', '2,04', '1,43\u20132,91', 'T\u0102NG 104%'],
     ['V\u1eafc-xin li\u1ec1u 1 (vs ch\u01b0a)', '0,97', '0,95\u20130,98', 'B\u1ea2O V\u1ec6 3%'],
     ['V\u1eafc-xin li\u1ec1u 2 (vs ch\u01b0a)', '0,95', '0,92\u20130,98', 'B\u1ea2O V\u1ec6 5%'],
     ['S\u1ee9c kh\u1ecfe k\u00e9m (vs \u1ed5n \u0111\u1ecbnh)', '1,71', '1,34\u20132,18', 'T\u0102NG 71%'],
     ['H\u1ecdc t\u1eadp k\u00e9m (vs \u1ed5n \u0111\u1ecbnh)', '1,44', '1,12\u20131,87', 'T\u0102NG 44%'],
     ['Quan h\u1ec7 XH k\u00e9m (vs \u1ed5n \u0111\u1ecbnh)', '1,54', '1,16\u20132,04', 'T\u0102NG 54%']],
    widths=[4.0, 1.5, 2.5, 3.5])
add_p(doc, 'COVID nhi\u1ec5m (PR=2,29) v\u00e0 ti\u1ebfp x\u00fac F0 (PR=2,04) l\u00e0 y\u1ebfu t\u1ed1 m\u1ea1nh nh\u1ea5t. V\u1eafc-xin c\u00f3 vai tr\u00f2 b\u1ea3o v\u1ec7 (PR=0,95\u20130,97). K\u00fd t\u00fac x\u00e1 + \u1edf m\u1ed9t m\u00ecnh t\u0103ng nguy c\u01a1 \u2014 c\u00f4 \u0111\u01a1n l\u00e0 y\u1ebfu t\u1ed1 n\u1ec1n? Ph\u00f9 h\u1ee3p Moon 2025 QT36 (c\u00f4 \u0111\u01a1n h\u1ea1ng 1 ML).', size=9, italic=True)

# 4. THAO LUAN
add_page_ref(doc, '4\u20135', 'Medicine', 'Vol. 102(16), 2023')
add_heading(doc, '4. TH\u1ea2O LU\u1eacN', 2)
add_p(doc, 'T\u1ea1i th\u1eddi \u0111i\u1ec3m NC, Vi\u1ec7t Nam \u0111ang gi\u1eefa \u0111\u1ee3t b\u00f9ng ph\u00e1t l\u1ea7n 4, trung b\u00ecnh 12.242 ca m\u1edbi/ng\u00e0y (t\u1ed5ng 435.265 ca; 10.749 t\u1eed vong), 62/63 t\u1ec9nh th\u00e0nh ghi nh\u1eadn ca COVID-19.')
add_p(doc, 'Lo \u00e2u 16,2% \u2014 th\u1ea5p h\u01a1n c\u00e1c n\u01b0\u1edbc kh\u00e1c (India, TQ Qu\u1ea3ng \u0110\u00f4ng, Th\u1ed5 Nh\u0129 K\u1ef3, Saudi Arabia). T\u01b0\u01a1ng t\u1ef1 NC kh\u00e1c t\u1ea1i VN cho th\u1ea5y ng\u01b0\u1eddi VN c\u00f3 m\u1ee9c lo \u00e2u th\u1ea5p h\u01a1n (so TQ, Iran) nh\u1edd s\u1ef1 ch\u1ec9 \u0111\u1ea1o c\u1ee7a Ch\u00ednh ph\u1ee7 VN v\u00e0 h\u00e0nh \u0111\u1ed9ng y t\u1ebf c\u1ed9ng \u0111\u1ed3ng ch\u1ee7 \u0111\u1ed9ng, hi\u1ec7u qu\u1ea3.')
add_p(doc, 'Gi\u1edbi t\u00ednh:', bold=True)
add_p(doc, 'N\u1eef c\u00f3 m\u1ee9c lo \u00e2u cao h\u01a1n nam (PR=0,66 cho nam). Ph\u00f9 h\u1ee3p NC tr\u01b0\u1edbc. N\u1eef c\u00f3 m\u1ee9c neuroticism/b\u1ea5t \u1ed5n c\u1ea3m x\u00fac cao h\u01a1n nam. Ph\u00f9 h\u1ee3p Islam 2025 (QT31: n\u1eef AOR=1,51 t\u1eeb 59 n\u01b0\u1edbc), Ireland 2024 (QT32: n\u1eef t\u0103ng nhanh h\u01a1n).')
add_p(doc, 'C\u01b0 tr\u00fa:', bold=True)
add_p(doc, 'K\u00fd t\u00fac x\u00e1 (PR=1,71) v\u00e0 \u1edf m\u1ed9t m\u00ecnh (PR=1,36) t\u0103ng nguy c\u01a1 lo \u00e2u so v\u1edbi s\u1ed1ng c\u00f9ng gia \u0111\u00ecnh. G\u1ee3i \u00fd: C\u00d4 \u0110\u01a0N v\u00e0 THI\u1ebeU H\u1ed6 TR\u1ee2 GIA \u0110\u00ccNH l\u00e0 y\u1ebfu t\u1ed1 n\u1ec1n. Ph\u00f9 h\u1ee3p Moon 2025 (QT36: c\u00f4 \u0111\u01a1n h\u1ea1ng 1 ML, F=16.646). Ph\u00f9 h\u1ee3p Ireland 2024 (QT32: OGA b\u1ea3o v\u1ec7).')
add_p(doc, 'COVID nhi\u1ec5m v\u00e0 v\u1eafc-xin:', bold=True)
add_p(doc, 'COVID nhi\u1ec5m (PR=2,29) v\u00e0 ti\u1ebfp x\u00fac F0 (PR=2,04) l\u00e0 y\u1ebfu t\u1ed1 m\u1ea1nh nh\u1ea5t. V\u1eafc-xin c\u00f3 vai tr\u00f2 b\u1ea3o v\u1ec7 c\u1ea3 s\u1ee9c kh\u1ecfe th\u1ec3 ch\u1ea5t v\u00e0 t\u00e2m th\u1ea7n (PR=0,95\u20130,97). Ph\u00f9 h\u1ee3p v\u1edbi Chen 2025 (meta COVID: COVID distress logOR=1,42).')
add_p(doc, 'H\u1ecdc t\u1eadp v\u00e0 quan h\u1ec7 XH:', bold=True)
add_p(doc, 'SV t\u1ef1 \u0111\u00e1nh gi\u00e1 h\u1ecdc t\u1eadp v\u00e0 quan h\u1ec7 XH "x\u1ea5u h\u01a1n" c\u00f3 lo \u00e2u cao h\u01a1n ~1,2\u20131,5 l\u1ea7n. NC tr\u01b0\u1edbc c\u0169ng cho th\u1ea5y ch\u1ea5t l\u01b0\u1ee3ng quan h\u1ec7 XH \u1ea3nh h\u01b0\u1edfng SKTT SV. Cam et al. (2021): n\u1eef gi\u1edbi v\u00e0 quan h\u1ec7 gia \u0111\u00ecnh k\u00e9m l\u00e0 y\u1ebfu t\u1ed1 nguy c\u01a1 cho PTSD, lo \u00e2u, tr\u1ea7m c\u1ea3m, stress. Ph\u00f9 h\u1ee3p Wen 2020 (QT08: h\u1ed7 tr\u1ee3 tr\u01b0\u1eddng OR=0,562 b\u1ea3o v\u1ec7), V\u0129nh L\u1ed9c 2024 (VN20: cha m\u1eb9 ly h\u00f4n t\u0103ng nguy c\u01a1).')

add_p(doc, 'H\u1ea1n ch\u1ebf:', bold=True)
add_p(doc, '(1) Thi\u1ebft k\u1ebf c\u1eaft ngang \u2014 kh\u00f4ng suy lu\u1eadn nh\u00e2n qu\u1ea3. (2) L\u1ea5y m\u1eabu t\u1ef1 ch\u1ecdn tr\u1ef1c tuy\u1ebfn \u2014 thi\u00ean l\u1ec7ch: ch\u1ec9 SV c\u00f3 internet v\u00e0 s\u1eb5n l\u00f2ng tham gia. (3) Ch\u1ec9 SV (\u226518 tu\u1ed5i) \u2014 kh\u00f4ng bao g\u1ed3m HS THCS/THPT (nh\u00f3m tu\u1ed5i quan tr\u1ecdng theo \u0111\u1ec1 t\u00e0i). (4) B\u1ed1i c\u1ea3nh COVID \u0111\u1eb7c bi\u1ec7t \u2014 k\u1ebft qu\u1ea3 c\u00f3 th\u1ec3 kh\u00f4ng kh\u00e1i qu\u00e1t cho giai \u0111o\u1ea1n b\u00ecnh th\u01b0\u1eddng.')

# 5. KET LUAN
add_heading(doc, '5. K\u1ebeT LU\u1eacN', 2)
add_p(doc, 'D\u1eef li\u1ec7u 5.730 SV VN trong l\u00e0n s\u00f3ng 4 COVID-19, cho th\u1ea5y lo \u00e2u 16,2% (GAD-7 \u226510), v\u1edbi COVID nhi\u1ec5m (PR=2,29), c\u01b0 tr\u00fa k\u00fd t\u00fac x\u00e1 (PR=1,71), s\u1ee9c kh\u1ecfe k\u00e9m (PR=1,71), quan h\u1ec7 XH k\u00e9m (PR=1,54), v\u00e0 h\u1ecdc t\u1eadp k\u00e9m (PR=1,44) l\u00e0 c\u00e1c y\u1ebfu t\u1ed1 li\u00ean quan, g\u1ee3i \u00fd r\u1eb1ng can thi\u1ec7p t\u00e2m l\u00fd c\u1ea7n nh\u1eafm v\u00e0o: (1) h\u1ed7 tr\u1ee3 SV c\u01b0 tr\u00fa xa gia \u0111\u00ecnh, (2) c\u1ea3i thi\u1ec7n m\u1ed1i quan h\u1ec7 x\u00e3 h\u1ed9i, (3) gi\u1ea3m t\u00e1c \u0111\u1ed9ng t\u00e2m l\u00fd COVID/kh\u1ee7ng ho\u1ea3ng s\u1ee9c kh\u1ecfe.')

# TLTK
add_heading(doc, 'T\u00c0I LI\u1ec6U THAM KH\u1ea2O', 2)
for ref in [
    'Nguyen, L.X., Dao, L.T.D., et al. (2023). Anxiety and associated factors among Vietnamese students during COVID-19 pandemic. Medicine, 102(16), e33559.',
    'Spitzer, R.L. et al. (2006). A brief measure for assessing GAD: the GAD-7. Arch Intern Med, 166(10), 1092\u20131097.',
    'Jiang, L. et al. (2021). Anxiety among students from 4 Asia-Pacific countries.',
    'Cam, N.N. et al. (2021). Risk factors for PTSD, anxiety, depression among Vietnamese during COVID-19.',
    '(Xem \u0111\u1ea7y \u0111\u1ee7 40 TLTK trong b\u00e0i g\u1ed1c)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('GAD-7', 'Generalized Anxiety Disorder 7-item Scale \u2014 Thang R\u1ed1i lo\u1ea1n Lo \u00e2u T\u1ed5ng qu\u00e1t 7 m\u1ee5c'),
    ('PR', 'Prevalence Ratio \u2014 T\u1ef7 s\u1ed1 T\u1ef7 l\u1ec7'),
    ('KTC', 'Kho\u1ea3ng tin c\u1eady (Confidence Interval)'),
    ('COVID-19', 'Coronavirus Disease 2019'),
    ('F0', 'Ng\u01b0\u1eddi nhi\u1ec5m COVID-19'),
    ('F1', 'Ng\u01b0\u1eddi ti\u1ebfp x\u00fac tr\u1ef1c ti\u1ebfp v\u1edbi F0'),
    ('SV', 'Sinh vi\u00ean'),
    ('SKTT', 'S\u1ee9c kh\u1ecfe t\u00e2m th\u1ea7n'),
    ('VASS', 'Vi\u1ec7n H\u00e0n l\u00e2m Khoa h\u1ecdc X\u00e3 h\u1ed9i Vi\u1ec7t Nam'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN \u0110I\u1ec2M PH\u1ea2N BI\u1ec6N / CRITICAL REVIEW')
add_red(doc, '\u0110i\u1ec3m m\u1ea1nh:', bold=True)
for s in [
    'Medicine Q1 (LWW) \u2014 t\u1ea1p ch\u00ed qu\u1ed1c t\u1ebf c\u00f3 PubMed. Open Access.',
    'M\u1eabu L\u1edaN NH\u1ea4T d\u00f9ng GAD-7 t\u1ea1i VN: n = 5.730 (so: Hoa 2024 VN01 n = 3.910, V-NAMHS VN02 n = 5.996 nh\u01b0ng d\u00f9ng DISC-5).',
    'GAD-7 \u03b1 = 0,92 \u2014 \u0111\u1ed9 tin c\u1eady c\u1ef1c cao. Ng\u01b0\u1ee1ng \u226510: sensitivity + specificity > 80%.',
    'Robust Poisson regression (PR) \u2014 ph\u00f9 h\u1ee3p h\u01a1n logistic (OR) cho d\u1eef li\u1ec7u t\u1ef7 l\u1ec7 > 10%. \u00cdt NC VN d\u00f9ng.',
    '16 t\u00e1c gi\u1ea3 t\u1eeb 9 c\u01a1 quan VN + \u00dac \u2014 h\u1ee3p t\u00e1c \u0111a ng\u00e0nh.',
    'B\u1ed1i c\u1ea3nh COVID l\u00e0n s\u00f3ng 4 \u2014 giai \u0111o\u1ea1n n\u1eb7ng nh\u1ea5t t\u1ea1i VN (12.242 ca m\u1edbi/ng\u00e0y).',
    'Ph\u00e1t hi\u1ec7n k\u00fd t\u00fac x\u00e1 PR=1,71 \u2014 g\u1ee3i \u00fd C\u00d4 \u0110\u01a0N l\u00e0 y\u1ebfu t\u1ed1 n\u1ec1n. Ph\u00f9 h\u1ee3p Moon 2025 QT36 (c\u00f4 \u0111\u01a1n h\u1ea1ng 1).',
]:
    add_red(doc, '\u2022 ' + s)

add_red(doc, 'H\u1ea1n ch\u1ebf chi ti\u1ebft:', bold=True)
for s in [
    'Ch\u1ec9 SINH VI\u00caN \u226518 tu\u1ed5i \u2014 KH\u00d4NG bao g\u1ed3m HS THCS/THPT (10\u201317 tu\u1ed5i). \u0110\u1ec1 t\u00e0i t\u1eadp trung c\u1ea5p 2+3 n\u00ean b\u00e0i n\u00e0y ch\u1ec9 l\u00e0 \u0111\u1ed1i chi\u1ebfu, kh\u00f4ng tr\u1ef1c ti\u1ebfp.',
    'L\u1ea5y m\u1eabu T\u1ef0 CH\u1eccN tr\u1ef1c tuy\u1ebfn \u2014 thi\u00ean l\u1ec7ch nghi\u00eam tr\u1ecdng: ch\u1ec9 SV c\u00f3 internet, s\u1ebb d\u00f9ng MXH, v\u00e0 s\u1eb5n l\u00f2ng tham gia. Lo\u1ea1i b\u1ecf SV kh\u00f4ng c\u00f3 internet/thi\u1ebft b\u1ecb.',
    'C\u1eaft ngang \u2014 kh\u00f4ng nh\u00e2n qu\u1ea3. So: VN21 (Tr\u1ea7n Th\u1ea3o Vi 2024) d\u00f9ng NC d\u1ecdc 3 n\u0103m.',
    'B\u1ed1i c\u1ea3nh COVID \u0111\u1eb7c bi\u1ec7t (l\u00e0n s\u00f3ng 4, gi\u00e3n c\u00e1ch x\u00e3 h\u1ed9i nghi\u00eam ng\u1eb7t) \u2014 k\u1ebft qu\u1ea3 c\u00f3 th\u1ec3 kh\u00f4ng kh\u00e1i qu\u00e1t. Ho\u00e0ng Trung H\u1ecdc 2025 (VN14): lo \u00e2u gi\u1ea3m t\u1eeb 41,5% xu\u1ed1ng 25,4% sau COVID.',
    'GAD-7 \u226510 \u2014 ch\u1ec9 \u0111o lo \u00e2u v\u1eeba\u2013n\u1eb7ng. B\u1ecf s\u00f3t lo \u00e2u nh\u1eb9 (5\u20139). Hoa 2024 (VN01) d\u00f9ng \u22655: 40,6% vs 16,2% \u2014 ch\u00eanh do ng\u01b0\u1ee1ng.',
    'Kh\u00f4ng \u0111o: \u00e1p l\u1ef1c h\u1ecdc t\u1eadp c\u1ee5 th\u1ec3 (ESSA), screen time/MXH, gi\u1ea5c ng\u1ee7, b\u1eaft n\u1ea1t, ACEs \u2014 y\u1ebfu t\u1ed1 quan tr\u1ecdng theo y v\u0103n.',
    'Kh\u00f4ng ph\u00e2n t\u1ea7ng theo V\u00d9NG MI\u1ec0N ho\u1eb7c \u0111\u1ed9 TU\u1ed4I chi ti\u1ebft.',
]:
    add_red(doc, '\u2022 ' + s)

add_red(doc, 'Kho\u1ea3ng tr\u1ed1ng NC / Gap:', bold=True)
for s in [
    'C\u1ea7n NC t\u01b0\u01a1ng t\u1ef1 cho HS THCS/THPT (10\u201317 tu\u1ed5i) \u2014 nh\u00f3m tu\u1ed5i tr\u1ecdng t\u00e2m c\u1ee7a \u0111\u1ec1 t\u00e0i. Lo \u00e2u \u1edf HS c\u00f3 th\u1ec3 kh\u00e1c SV (kh\u00f4ng s\u1ed1ng xa gia \u0111\u00ecnh nh\u01b0ng ch\u1ecbu \u00e1p l\u1ef1c thi c\u1eed).',
    'So s\u00e1nh GAD-7 \u22655 vs \u226510 tr\u00ean c\u00f9ng m\u1eabu \u2014 x\u00e1c \u0111\u1ecbnh ng\u01b0\u1ee1ng ph\u00f9 h\u1ee3p cho VTN VN. Hoa 2024 (VN01) d\u00f9ng \u22655: 40,6%; b\u00e0i n\u00e0y d\u00f9ng \u226510: 16,2%.',
    'NC h\u1eadu COVID: lo \u00e2u SV sau COVID c\u00f3 GI\u1ea2M kh\u00f4ng? Ho\u00e0ng Trung H\u1ecdc (VN14): VTN gi\u1ea3m t\u1eeb 41,5% xu\u1ed1ng 25,4%.',
    'K\u00fd t\u00fac x\u00e1 = y\u1ebfu t\u1ed1 nguy c\u01a1 \u2014 c\u1ea7n NC c\u00f4 \u0111\u01a1n + thi\u1ebfu h\u1ed7 tr\u1ee3 XH nh\u01b0 y\u1ebfu t\u1ed1 trung gian. Moon 2025 (QT36): c\u00f4 \u0111\u01a1n h\u1ea1ng 1 ML. Ireland 2024 (QT32): OGA b\u1ea3o v\u1ec7.',
    'Th\u00eam bi\u1ebfn: screen time (Norway QT21), gi\u1ea5c ng\u1ee7 (Zhu QT05: AOR=13,71), \u00e1p l\u1ef1c HT (Wen QT08: OR=11,58) v\u00e0o m\u00f4 h\u00ecnh.',
]:
    add_red(doc, '\u2022 ' + s)

# SAVE
outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '39_VN_COVID_Medicine_2023.docx')
doc.save(outpath)

# Verify
import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['5.730', '16,2', '0,66', '1,71', '2,29', '0,97', '0,95', '1,44', '1,54', 'GAD-7', '0,92']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
