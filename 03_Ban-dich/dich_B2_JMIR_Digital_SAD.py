# -*- coding: utf-8 -*-
"""Dịch B2 — JMIR Digital SAD 2025 — Walder et al."""
import sys, os
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

doc = create_doc()
add_p(doc, 'Link: https://doi.org/10.2196/preprints.67067', size=10)
add_p(doc, 'JMIR: https://www.jmir.org/2025/1/e67067', size=10)

add_heading(doc, 'Can thiệp sức khỏe tâm thần kỹ thuật số cho phòng ngừa và điều trị rối loạn lo âu xã hội ở trẻ em, thanh thiếu niên và thanh niên: Tổng quan hệ thống và phân tích tổng hợp các thử nghiệm ngẫu nhiên có đối chứng', 1)
h = doc.add_paragraph()
r = h.add_run('Digital Mental Health Interventions for Prevention and Treatment of Social Anxiety Disorder in Children, Adolescents and Young Adults: Systematic Review and Meta-Analysis of Randomised Controlled Trials')
r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.italic = True

add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
add_info_table(doc, [
    ('Tiêu đề gốc', 'Digital Mental Health Interventions for Prevention and Treatment of Social Anxiety Disorder in Children, Adolescents and Young Adults: SR and MA of RCTs'),
    ('Tác giả', 'Noemi Walder PhD, Alessja Frey MSc, Thomas Berger PhD, Stefanie Julia Schmidt PhD'),
    ('Cơ quan', 'Division of Clinical Child and Adolescent Psychology + Clinical Psychology and Psychotherapy, ĐH Bern, Thụy Sĩ'),
    ('Tạp chí', 'Journal of Medical Internet Research (JMIR) Q1 IF ≈ 7,4'),
    ('Xuất bản', '2025, 45 trang (preprint accepted)'),
    ('DOI', '10.2196/preprints.67067'),
    ('PROSPERO', 'CRD42023424181'),
    ('Loại NC', 'Tổng quan hệ thống + Phân tích tổng hợp (Random-effects)'),
    ('Mẫu', '22 NC (21 meta-analysis), RCT, trẻ/VTN/thanh niên (TB <25 tuổi)'),
])
add_page_ref(doc, '1–45', 'JMIR', '2025')

# TOM TAT
add_heading(doc, 'TÓM TẮT', 2)
add_p(doc, 'Bối cảnh: Rối loạn lo âu xã hội (SAD) ảnh hưởng đáng kể đến chức năng xã hội và học tập của giới trẻ, nhấn mạnh nhu cầu điều trị tiếp cận được và hiệu quả như can thiệp sức khỏe tâm thần kỹ thuật số (DMHI — Digital Mental Health Interventions).')

p = doc.add_paragraph()
r = p.add_run('Phương pháp: Tìm kiếm RCT đánh giá DMHI cho lo âu xã hội ở người trẻ (TB <25 tuổi). 2 tác giả độc lập sàng lọc, trích xuất, đánh giá thiên lệch. k=22 tổng quan, k=21 meta-analysis.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

p = doc.add_paragraph()
r = p.add_run('Kết quả: DMHI hiệu quả đáng kể: g = 0,508 (KTC 95%: 0,308–0,707) so với BẤT KỲ đối chứng nào. Hiệu quả mạnh hơn khi: (1) so với WL (g = 0,576), (2) dựa trên CBT (g = 0,610), (3) có thành phần SAD riêng (g = 0,878 — LỚN NHẤT), (4) có hướng dẫn người (g = 0,825). Tuổi và sự tham gia cha mẹ KHÔNG ảnh hưởng đáng kể. Khi xét thiên lệch xuất bản: hiệu quả vẫn đáng kể (g = 0,506).')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r.font.color.rgb = RGBColor(0xCC, 0, 0)

add_p(doc, 'Kết luận: Meta-analysis ủng hộ hiệu quả DMHI cho lo âu xã hội so với không can thiệp, và lợi ích của hướng dẫn + can thiệp thiết kế riêng cho SAD. Tuy nhiên, nhấn mạnh thiếu sót phương pháp và chuẩn báo cáo không đồng nhất.')

# DANH GIA NHANH
add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
for b in [
    'JMIR Q1 IF ≈ 7,4 — tạp chí hàng đầu về sức khỏe số.',
    'META-ANALYSIS 21 RCT — toàn diện nhất về DMHI cho SAD ở trẻ/VTN/thanh niên.',
    'DMHI g = 0,508 (TRUNG BÌNH) — hiệu quả có ý nghĩa lâm sàng.',
    'SAD-specific g = 0,878 — CAN THIỆP RIÊNG CHO SAD hiệu quả GẤP ĐÔI can thiệp chung.',
    'Có hướng dẫn người g = 0,825 — HƯỚNG DẪN quan trọng (không chỉ tự học).',
    'CBT-based g = 0,610 — CBT vẫn là nền tảng tốt nhất.',
    'Tuổi KHÔNG ảnh hưởng — áp dụng được cả trẻ em và VTN.',
    'PROSPERO CRD42023424181.',
]:
    add_p(doc, f'• {b}')
add_p(doc, 'Hạn chế:', bold=True)
for b in [
    'Preprint (accepted) — chưa phải bản cuối cùng.',
    'Đa số RCT từ phương Tây — ít châu Á/LMIC.',
    'Hiệu quả CÓ THỂ MẤT sau 12 tháng (không đáng kể follow-up dài).',
    'SAD riêng — không áp dụng cho GAD/lo âu chung.',
    'Chất lượng NC: heterogeneous, nhiều NC risk of bias.',
    'Không đánh giá chi phí-hiệu quả.',
]:
    add_p(doc, f'• {b}')

# KET QUA CHI TIET
add_heading(doc, 'KẾT QUẢ CHI TIẾT', 2)

add_heading(doc, 'Bảng 1. Hiệu quả DMHI — Phân tích phân nhóm', 3)
add_table(doc,
    ['Phân nhóm', 'g (Hedges)', 'KTC 95%', 'Ý nghĩa'],
    [['DMHI vs BẤT KỲ đối chứng', '0,508', '0,308–0,707', 'TRUNG BÌNH — đáng kể'],
     ['DMHI vs WL (danh sách chờ)', '0,576', '0,343–0,809', 'Mạnh hơn vs WL'],
     ['DMHI dựa trên CBT', '0,610', '0,361–0,859', 'CBT nền tảng TỐT NHẤT'],
     ['DMHI có thành phần SAD riêng', '0,878', '0,469–1,278', 'LỚN NHẤT — SAD-specific vượt trội'],
     ['DMHI có hướng dẫn người', '0,825', '0,425–1,224', 'Hướng dẫn QUAN TRỌNG'],
     ['DMHI không hướng dẫn', '~0,3–0,4', '—', 'Nhỏ — kém hiệu quả hơn'],
     ['Sau xét thiên lệch XB', '0,506', '0,308–0,707', 'Vẫn đáng kể — robust']],
    widths=[4.5, 2.0, 2.5, 4.0])

add_heading(doc, 'Bảng 2. So sánh với NC can thiệp khác trong đề tài', 3)
add_table(doc,
    ['NC', 'Loại', 'Kết quả', 'So sánh'],
    [['B2 JMIR (bài này)', 'DMHI SAD MA', 'g = 0,508 tổng; SAD-specific 0,878', 'Digital — trực tuyến'],
     ['B9 NMA SAD (Xian 2024)', 'NMA 30 RCT SAD', 'iCBT SUCRA 71,2% hạng 1', 'iCBT = DMHI — PHÙ HỢP'],
     ['QT29 BMC NMA (Li 2025)', 'NMA 30 RCT lo âu chung', 'ACT 0,69; CBT 0,66', 'Lo âu chung (không SAD)'],
     ['QT33 JAMA RCT', 'RCT giảm screen time', 'Cohen d = 0,53', 'Giảm ST — khác CBT'],
     ['B8 Sri Lanka RCT', 'Cluster RCT CBT trường', 'Lo âu giảm β = −0,096', 'Mặt đối mặt — trường'],
     ['QT28 AJP (Zugman)', 'Tổng quan CBT+SSRI', '80,7% đáp ứng (CAMS)', 'CBT+thuốc — lâm sàng']],
    widths=[3.0, 2.5, 3.5, 3.5])
add_p(doc, 'B2 (DMHI SAD g=0,508) + B9 (iCBT SAD SUCRA 71,2%) → bằng chứng NHẤT QUÁN: can thiệp KỸ THUẬT SỐ hiệu quả cho SAD ở trẻ/VTN. Đặc biệt khi có HƯỚNG DẪN (g=0,825) và THIẾT KẾ RIÊNG SAD (g=0,878).', size=9, italic=True)

# THAO LUAN
add_heading(doc, 'THẢO LUẬN VÀ KẾT LUẬN', 2)
add_p(doc, 'DMHI hiệu quả cho SAD ở trẻ/VTN — phát hiện quan trọng. 3 yếu tố tăng hiệu quả: (1) dựa trên CBT (g=0,610), (2) thiết kế riêng cho SAD (g=0,878 — GẤP ĐÔI can thiệp chung), (3) có hướng dẫn người (g=0,825). Gợi ý: app SKTT cần CÓ NGƯỜI HƯỚNG DẪN (therapist/GV/mentor), không chỉ tự học.')
add_p(doc, 'Kết hợp B2 + B9: iCBT (internet CBT) là can thiệp tối ưu cho SAD ở VTN — cả meta-analysis (g=0,508–0,878) và NMA (SUCRA 71,2%) đều xác nhận. Hàm ý VN: phát triển app iCBT tiếng Việt cho SAD, có GV/mentor hướng dẫn, thiết kế riêng cho lo âu XH (không phải lo âu chung).')
add_p(doc, 'Tuổi không ảnh hưởng → áp dụng cả THCS và THPT. Cha mẹ tham gia không ảnh hưởng → khác với B8 Sri Lanka (phiên cha mẹ). Có thể do DMHI = tự học, cha mẹ ít ảnh hưởng.')
add_p(doc, 'Hạn chế: hiệu quả có thể MẤT sau 12 tháng → cần duy trì can thiệp. Đa số NC phương Tây → cần thử ở châu Á/VN (B7 CA-CBT ĐNA: cần thích ứng văn hóa).')

# TLTK
add_heading(doc, 'TÀI LIỆU THAM KHẢO', 2)
for ref in [
    'Walder, N., Frey, A., Berger, T. & Schmidt, S.J. (2025). Digital MH Interventions for SAD in Children, Adolescents and Young Adults: SR and MA of RCTs. JMIR, e67067.',
    '(Xem đầy đủ 100+ TLTK trong bài gốc 45 trang)',
]:
    add_p(doc, ref, size=10)

# VIET TAT
add_abbreviation_table(doc, [
    ('DMHI', 'Digital Mental Health Interventions — Can thiệp SKTT Kỹ thuật số'),
    ('SAD', 'Social Anxiety Disorder — Rối loạn Lo âu Xã hội'),
    ('iCBT', 'Internet-delivered CBT — CBT qua Internet'),
    ('CBT', 'Cognitive Behavioral Therapy'),
    ('WL', 'Waiting List — Danh sách Chờ'),
    ('g', 'Hedges g — Kích thước hiệu ứng (chuẩn hóa)'),
    ('RCT', 'Randomized Controlled Trial'),
    ('PROSPERO', 'Prospective Register of Systematic Reviews'),
    ('JMIR', 'Journal of Medical Internet Research'),
])

# PHAN BIEN
add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
add_red(doc, 'Điểm mạnh:', bold=True)
for s in [
    'JMIR Q1 IF ≈ 7,4. META-ANALYSIS 21 RCT — toàn diện nhất về DMHI cho SAD.',
    'Phân tích phân nhóm chi tiết: CBT-based, SAD-specific, hướng dẫn, WL vs active. PROSPERO.',
    'SAD-specific g = 0,878 — phát hiện ĐỘT PHÁ: can thiệp RIÊNG SAD vượt trội chung. Phù hợp B9 NMA SAD.',
    'Hướng dẫn người g = 0,825 — xác nhận cần NGƯỜI HỖ TRỢ (không chỉ app tự học).',
    'Thiên lệch XB xét → vẫn đáng kể (g = 0,506) — robust.',
    'Tuổi không ảnh hưởng → áp dụng THCS + THPT + đại học.',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Hạn chế chi tiết:', bold=True)
for s in [
    'Preprint accepted — chưa phải bản cuối. Có thể có chỉnh sửa.',
    'Đa số RCT phương Tây — ít châu Á/LMIC. B7 CA-CBT ĐNA: chỉ 7 NC toàn ĐNA.',
    'Hiệu quả có thể MẤT sau 12 tháng — cần duy trì liên tục.',
    'Heterogeneous — NC khác nhau về thiết kế, mẫu, đo lường.',
    'Không đánh giá chi phí-hiệu quả — quan trọng cho LMIC.',
    'SAD riêng — không áp dụng GAD (cần B9 NMA cho lo âu chung).',
]:
    add_red(doc, f'• {s}')
add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
for s in [
    'VN chưa có DMHI/app nào cho SAD ở trẻ/VTN — GAP cực lớn. Jefferies QT35: VN SAD = 30,7%.',
    'Phát triển app iCBT tiếng Việt cho SAD + hướng dẫn GV/mentor — kết hợp B2 (DMHI) + B7 (CA-CBT ĐNA) + B8 (GV cung cấp).',
    'RCT so sánh: app iCBT SAD VN vs gCBT mặt đối mặt — cái nào hiệu quả + chi phí-hiệu quả hơn?',
    'Duy trì hiệu quả: theo dõi 6–12 tháng. Booster sessions (phiên nhắc lại) để duy trì?',
    'Thích ứng văn hóa VN: tình huống lo âu XH VN (nói trước lớp, gặp người có thẩm quyền — QT35 xếp hạng 3/17).',
]:
    add_red(doc, f'• {s}')

outpath = os.path.join(os.path.dirname(os.path.abspath(__file__)), '44_JMIR_Digital_SAD_2025.docx')
doc.save(outpath)

import docx as dx
d = dx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            t += ' ' + cell.text
checks = ['0,508', '0,576', '0,610', '0,878', '0,825', '0,506', '21 RCT', 'DMHI', 'Bayesian', 'PROSPERO']
ok = sum(1 for c in checks if c in t)
print(f'Saved: {outpath}')
print(f'  Chars: {len(t)}, Tables: {len(d.tables)}, Paras: {len([p for p in d.paragraphs if p.text.strip()])}')
print(f'  Numbers verified: {ok}/{len(checks)}')
