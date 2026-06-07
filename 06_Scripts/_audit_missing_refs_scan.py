"""Read-only scan of Tom-tat-tung-bai/ summary docx for missing References section."""
import os
import re
import json
from pathlib import Path
from docx import Document

ROOT = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au/Tom-tat-tung-bai")

HEADING_PATTERNS = [
    r"tham\s*kh[aả]o",
    r"references?\b",
    r"t[aà]i\s*li[eệ]u\s*tham\s*kh[aả]o",
    r"\btltk\b",
    r"bibliography",
]
HEADING_RE = re.compile("|".join(HEADING_PATTERNS), re.IGNORECASE)
# Heading line: optionally numbered "7. " or "VII. " or "#" prefixes
HEADING_LINE_RE = re.compile(
    r"^\s*(?:[#\d]+[\.\):]?\s+|[IVX]+\.\s+)?(?:" + "|".join(HEADING_PATTERNS) + r")",
    re.IGNORECASE,
)

DOI_RE = re.compile(r"\b10\.\d{4,9}/[-._;()/:A-Z0-9]+", re.IGNORECASE)
PMID_RE = re.compile(r"\bPMID[:\s]*\d+", re.IGNORECASE)
URL_RE = re.compile(r"https?://[^\s)]+", re.IGNORECASE)
# Simple inline citation marker like (Author, 2023) or [12]
CIT_PAREN_RE = re.compile(r"\([A-ZĐ][^()]{0,40},\s*(?:19|20)\d{2}[a-z]?\)")
CIT_BRACKET_RE = re.compile(r"\[\d{1,3}(?:[,\-]\s*\d{1,3})*\]")

# Skip patterns
SKIP_NAMES = {"BANG_LOI_DA_SUA.docx", "BANG_GIOI_TINH_LIEN_BAI.docx",
              "00_Mẫu tóm tắt bài 1_FIXED_27052026.docx"}


def classify(docx_path: Path):
    try:
        doc = Document(str(docx_path))
    except Exception as e:
        return {"file": docx_path.name, "error": str(e)}

    paragraphs = [(p.style.name if p.style else "", p.text or "") for p in doc.paragraphs]
    full_text = "\n".join(t for _, t in paragraphs)

    # Find heading match index
    heading_idx = None
    heading_kind = None
    for i, (style, text) in enumerate(paragraphs):
        t = text.strip()
        if not t:
            continue
        # Heading detection: either style is Heading*, or text matches and is short (<= 80 chars)
        is_heading_style = style.lower().startswith("heading") or style.lower().startswith("title")
        m = HEADING_RE.search(t)
        if m and (is_heading_style or len(t) <= 80):
            # Heuristic: avoid matching inline mentions like "trong phần Tham khảo của bài..."
            # Accept if heading-like (style) OR matches at start (incl. numbered prefix)
            if is_heading_style or HEADING_LINE_RE.match(t):
                heading_idx = i
                heading_kind = t
                break

    has_section_content = False
    if heading_idx is not None:
        # Check next 50 paragraphs for non-empty content
        for j in range(heading_idx + 1, min(heading_idx + 80, len(paragraphs))):
            t = paragraphs[j][1].strip()
            if t and len(t) > 10:
                has_section_content = True
                break

    # Also scan tables for ref-like content
    table_text_parts = []
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                table_text_parts.append(cell.text or "")
    table_text = "\n".join(table_text_parts)

    all_text = full_text + "\n" + table_text

    n_doi = len(DOI_RE.findall(all_text))
    n_pmid = len(PMID_RE.findall(all_text))
    n_url = len(URL_RE.findall(all_text))
    n_cit_paren = len(CIT_PAREN_RE.findall(all_text))
    n_cit_bracket = len(CIT_BRACKET_RE.findall(all_text))

    has_inline_refs = (n_doi + n_pmid + n_url) > 0

    if has_section_content:
        cls = "HAS_REF_SECTION"
    elif has_inline_refs:
        cls = "HAS_INLINE_REFS"
    else:
        cls = "NO_REFS"

    return {
        "file": docx_path.name,
        "classification": cls,
        "heading_found": heading_kind,
        "section_has_content": has_section_content,
        "n_doi": n_doi,
        "n_pmid": n_pmid,
        "n_url": n_url,
        "n_cit_paren": n_cit_paren,
        "n_cit_bracket": n_cit_bracket,
        "n_paragraphs": len(paragraphs),
        "n_tables": len(doc.tables),
    }


def main():
    files = sorted([p for p in ROOT.iterdir()
                    if p.is_file() and p.suffix.lower() == ".docx"
                    and p.name not in SKIP_NAMES])
    results = []
    for f in files:
        r = classify(f)
        results.append(r)

    # Print summary
    counts = {"HAS_REF_SECTION": 0, "HAS_INLINE_REFS": 0, "NO_REFS": 0, "error": 0}
    for r in results:
        if "error" in r:
            counts["error"] += 1
        else:
            counts[r["classification"]] += 1

    print(json.dumps({"total": len(results), "counts": counts}, ensure_ascii=False))
    # Dump full results to JSON
    out = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au/06_Scripts/_audit_missing_refs_results.json")
    out.write_text(json.dumps(results, ensure_ascii=False, indent=2), encoding="utf-8")
    print(f"Wrote {out}")


if __name__ == "__main__":
    main()
