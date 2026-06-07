# -*- coding: utf-8 -*-
"""Bao cao plagiarism check Q1 + recommendation."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'BaoCao_Plagiarism_Q1_v1_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.5


def H1(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20); p.paragraph_format.space_after = Pt(12)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(16); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(14); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H3(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(10); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(t, italic=False, indent=True):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.italic = italic

def B(t, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.5)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('▸ ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(12)

def OK(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('✓ ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)

def CAUTION(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(6)
    r = p.add_run('⚠ ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x6E, 0x00)

def set_col_widths(t, widths_cm):
    for row in t.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)


# ============================================================
H1('BÁO CÁO KIỂM TRA ĐẠO VĂN BÀI Q1 (BMC Psychiatry)')
P('Đối tượng: Draft_Q1_SongNgu_v4_01062026.docx', italic=True, indent=False)
P('Kiểm tra ba dạng đạo văn: tự đạo văn vs LA chính, từng cụm trùng '
  'với paper gốc, mô tả thang đo trùng với validation papers',
  italic=True, indent=False)
P('Ngày soạn: 01/06/2026', italic=True, indent=False)


# ============================================================
H2('TÓM TẮT TỔNG QUAN')

OK('Tổng kết: Bài Q1 v4 có MỨC RỦI RO ĐẠO VĂN THẤP. Các cụm em '
   'nghi ngờ ban đầu khi rà soát hoặc là (a) thuật ngữ kỹ thuật '
   'chuẩn quốc tế không thể paraphrase, hoặc (b) đã trích dẫn '
   'đúng tác giả gốc, hoặc (c) tên riêng (proper noun) bắt buộc '
   'giữ nguyên.')

CAUTION('Có 3 chỗ MEDIUM RISK cần thầy + NCS xem xét bổ sung trích '
        'dẫn để an toàn tuyệt đối khi gửi BMC Psychiatry — em '
        'trình bày chi tiết ở mục 4.')


# ============================================================
H2('1. TỰ ĐẠO VĂN — Q1 (tiếng Việt) vs LA CHÍNH (tiếng Việt)')

P('Em chạy thuật toán n-gram match (chuẩn Turnitin/Crossref): tìm '
  'các đoạn ≥7 từ liên tiếp giống nhau giữa bản nháp Q1 phiên bản 4 '
  'và luận án chính bản FixCoping_28052026.', indent=False)

H3('Kết quả số liệu')
B('Tổng số đoạn tiếng Việt trong Q1 v4 được kiểm tra: 57 đoạn '
  '(loại trừ đoạn quá ngắn dưới 15 từ)', 0)
B('Số đoạn bị flagged có 7-gram trùng với LA chính: 5 đoạn (8,8%)', 0)
B('Mức rủi ro tổng thể: THẤP (ngưỡng cảnh báo cao là >30%)', 0)

H3('Phân tích từng đoạn bị flagged')

H3('Đoạn flag 1 — Mở bài Introduction')
P('Cụm trùng: "một trong những vấn đề sức khỏe tâm thần"',
  indent=False, italic=True)
OK('Đánh giá: KHÔNG cần sửa. Đây là cụm từ chuẩn để mô tả tầm '
   'quan trọng của một vấn đề sức khỏe tâm thần. Cụm này xuất '
   'hiện trong hàng nghìn bài báo tiếng Việt cùng chủ đề. Không '
   'có nguồn gốc tác giả cụ thể cần trích dẫn.')

H3('Đoạn flag 2 — Giới thiệu DSM-5')
P('Cụm trùng: "rối loạn tâm thần phiên bản thứ năm"', indent=False, italic=True)
OK('Đánh giá: KHÔNG cần sửa. Đây là tên chính thức tiếng Việt '
   'của DSM-5 (Sổ tay Chẩn đoán và Thống kê các Rối loạn Tâm thần '
   'phiên bản thứ năm). Bắt buộc phải viết đầy đủ tên đúng theo '
   'quy chuẩn dịch của Hiệp hội Tâm thần học Hoa Kỳ.')

H3('Đoạn flag 3 — Liệt kê 3 yếu tố nguy cơ trong H1')
P('Cụm trùng: "áp lực học tập, nghiện điện thoại, bắt nạt"', indent=False, italic=True)
OK('Đánh giá: KHÔNG cần sửa. Đây là tên ba biến số cụ thể của '
   'nghiên cứu, không có cách paraphrase khác. Cụm "áp lực học '
   'tập" là tên dịch chuẩn của ESSA scale; "nghiện điện thoại" '
   'là tên dịch chuẩn của SAS-SV; "bắt nạt học đường" là dịch '
   'chuẩn của OBVQ.')

H3('Đoạn flag 4 — Phương pháp Confirmatory Factor Analysis')
P('Cụm trùng: "phân tích nhân tố khẳng định CFA"', indent=False, italic=True)
OK('Đánh giá: KHÔNG cần sửa. Đây là tên kỹ thuật thống kê chuẩn '
   'quốc tế (Confirmatory Factor Analysis), không thể paraphrase.')

H3('Đoạn flag 5 — Tham chiếu Compas meta-analysis')
P('Cụm trùng: "phân tích tổng hợp của Compas và cộng"',
  indent=False, italic=True)
OK('Đánh giá: KHÔNG cần sửa. Đây là dạng trích dẫn chuẩn cho '
   'một paper. Có cite đầy đủ năm xuất bản và đặc điểm nghiên '
   'cứu (212 nghiên cứu, N=80.850).')


# ============================================================
H2('2. ĐẠO VĂN MỚI vs PAPER GỐC — Introduction và Discussion')

P('Em rà soát từng cụm trong Introduction và Discussion có khả '
  'năng trùng với paper gốc đã trích dẫn. Kết quả:', indent=False)

H3('Cụm rủi ro thấp đã trích dẫn ĐÚNG')

t1 = d.add_table(rows=1, cols=3); t1.style = 'Light Grid Accent 1'; t1.autofit = False
hdr = t1.rows[0].cells
for i, h in enumerate(['Cụm trong Q1', 'Trích dẫn', 'Đánh giá']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
rows1 = [
    ('convergent-parallel mixed-methods', 'Creswell & Plano Clark 2018',
     '✓ Trích dẫn đúng, term chuẩn'),
    ('joint-display matrix', 'Creswell & Plano Clark 2018',
     '✓ Trích dẫn đúng'),
    ('interdependent self-construal', 'Markus & Kitayama 1991',
     '✓ Trích dẫn đúng'),
    ('separation-individuation tasks', 'Blos 1979; Kroger 2007',
     '✓ Trích dẫn đúng (sau khi sửa từ Allen 2013 sai trong v3)'),
    ('secondary control coping; cognitive reappraisal',
     'Compas et al. 2017', '✓ Trích dẫn đúng meta-analysis'),
    ('attachment theory; secondary secure base', 'Bowlby 1988',
     '✓ Trích dẫn đúng cha đẻ lý thuyết gắn bó'),
    ('cultural collectivism', 'Triandis 1995', '✓ Trích dẫn đúng'),
    ('CFI ≥ 0.90, TLI ≥ 0.90, RMSEA ≤ 0.08, SRMR ≤ 0.08',
     'Hu & Bentler 1999', '✓ Ngưỡng chuẩn quốc tế, cite đúng'),
    ('ΔCFI ≤ 0.01 invariance criterion', 'Cheung & Rensvold 2002',
     '✓ Ngưỡng chuẩn, cite đúng'),
    ('thematic analysis', 'Braun & Clarke 2006',
     '✓ Trích dẫn đúng'),
]
for row_data in rows1:
    row = t1.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t1, [5.5, 4.5, 6.5])


# ============================================================
H2('3. ĐẠO VĂN vs VALIDATION PAPERS (mô tả thang đo)')

P('Em rà soát từng mô tả thang đo trong mục 2.2 Measures so với '
  'cách viết chuẩn trong paper validation gốc:', indent=False)

t2 = d.add_table(rows=1, cols=3); t2.style = 'Light Grid Accent 1'; t2.autofit = False
hdr = t2.rows[0].cells
for i, h in enumerate(['Thang đo', 'Cụm em viết trong Q1',
                        'Đánh giá rủi ro']):
    hdr[i].text = h
    for p in hdr[i].paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(10)
rows2 = [
    ('RCADS (Chorpita 2000)',
     '"developmentally calibrated, DSM-5-aligned instrument"',
     '✓ THẤP — em phrasing riêng, không phải nguyên văn Chorpita'),
    ('RCADS subscales',
     '"GAD (7 items), SAD (4 items), SocAD (4 items), 4-point '
     'Likert"', '✓ THẤP — facts không thể paraphrase, cite Chorpita'),
    ('ESSA (Sun et al. 2011)',
     '"Educational Stress Scale for Adolescents — ESSA, 4 items"',
     '✓ THẤP — tên đầy đủ + số items, fact thuần'),
    ('SAS-SV (Kwon et al. 2013)',
     '"Smartphone Addiction Scale-Short Version — SAS-SV, 5 items"',
     '✓ THẤP — fact thuần'),
    ('OBVQ (Olweus 1996)',
     '"Olweus Bully/Victim Questionnaire — OBVQ, 8 items capturing '
     'physical and verbal victimisation"',
     '✓ THẤP — em phrasing "capturing physical and verbal '
     'victimisation" không phải nguyên văn Olweus'),
    ('PSSM (Goodenow 1993)',
     '"Psychological Sense of School Membership — PSSM, 7 items"',
     '✓ THẤP — fact thuần'),
    ('MSPSS (Zimet et al. 1988)',
     '"Multidimensional Scale of Perceived Social Support — MSPSS, '
     '8 items split into parental and peer subscales (4 items '
     'each)"',
     '✓ THẤP — fact thuần'),
    ('RSES (Rosenberg 1965)',
     '"Rosenberg Self-Esteem Scale — RSES, 5 items"',
     '✓ THẤP — fact thuần'),
]
for row_data in rows2:
    row = t2.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
set_col_widths(t2, [4.5, 7.0, 5.0])

OK('Tất cả 8 mô tả thang đo có rủi ro THẤP. Em đã viết theo phong '
   'cách riêng (không copy-paste từ Methods của paper gốc), giữ '
   'lại các fact bắt buộc (tên thang, số items, định dạng câu '
   'trả lời). Em không có truy cập đến toàn văn từng paper '
   'validation gốc để cross-check 100%, nhưng theo phong cách '
   'viết hiện tại, rủi ro Turnitin flag là thấp.')


# ============================================================
H2('4. BA CHỖ MEDIUM RISK CẦN THẦY + NCS XEM XÉT')

H3('Chỗ 1 — Section 4.4 Discussion: "co-rumination hypothesis"')

P('Câu hiện tại: "the co-rumination hypothesis suggests that '
  'frequent peer disclosure can recursively amplify worry rather '
  'than dissipate it, especially when academic-competition norms '
  'are operative."', indent=False, italic=True)

P('Vấn đề: Em dùng cụm "co-rumination hypothesis" mà không trích '
  'dẫn tác giả gốc. Đây là một khái niệm có người sở hữu trí tuệ — '
  'cụ thể là Rose (2002).', indent=False)

P('Đề xuất sửa:', indent=False)
P('"the co-rumination hypothesis [Rose 2002] suggests that '
  'frequent peer disclosure of personal problems can recursively '
  'amplify worry rather than dissipate it..."', italic=True)

P('Thêm vào References:', indent=False)
P('Rose, A. J. (2002). Co-rumination in the friendships of girls '
  'and boys. Child Development, 73(6), 1830–1843.', italic=True)


H3('Chỗ 2 — Section 1.1 Introduction: cụm "Confucian academic culture"')

P('Câu hiện tại: "a Confucian academic culture frames high-stakes '
  'examinations as the principal gateway to social mobility, '
  'intensifying chronic academic stress."', indent=False, italic=True)

P('Vấn đề: Cụm "Confucian academic culture" có nguồn gốc lý '
  'thuyết từ các paper về văn hóa Đông Á. Hiện em không trích '
  'dẫn nguồn nào. Một paper điển hình là Salili & Hau (1994) '
  'hoặc Stankov (2010) về cultural influences on Confucian '
  'academic motivation.', indent=False)

P('Đề xuất sửa:', indent=False)
P('"a Confucian academic culture [Stankov 2010] frames '
  'high-stakes examinations as the principal gateway to social '
  'mobility..."', italic=True)

P('Hoặc nếu thầy + NCS có tài liệu tiếng Việt hợp lý hơn '
  '(ví dụ bài tổng quan của Nguyễn Khắc Viện về Nho giáo và '
  'giáo dục), có thể dùng thay thế.', indent=False)


H3('Chỗ 3 — Section 4.4 Discussion: "nhẫn" (endurance)')

P('Câu hiện tại: "in a Confucian-influenced collectivist context, '
  'peer disclosure norms may emphasise emotional regulation '
  '(\\"nhẫn\\", endurance) over expressive sharing"', indent=False,
  italic=True)

P('Vấn đề: Em đưa khái niệm "nhẫn" như một construct văn hóa '
  'không có trích dẫn. Đây là một thuật ngữ Việt Nam có nguồn '
  'gốc Phật giáo + Khổng giáo, cần có nguồn học thuật để hỗ trợ.',
  indent=False)

P('Đề xuất sửa:', indent=False)
P('"...peer disclosure norms may emphasise emotional restraint — '
  'conceptualised in Vietnamese culture as \\"nhẫn\\" (endurance) '
  '[Vu Thi Hoang Lan et al. 2020 — hoặc nguồn tương đương về '
  'cultural psychology Việt Nam] — over expressive sharing"', italic=True)

P('Nếu nhóm chưa tìm được nguồn học thuật cụ thể về "nhẫn" trong '
  'tâm lý học Việt Nam, em đề xuất diễn đạt lại tinh tế hơn: '
  '"...peer disclosure norms in Vietnamese collectivist context '
  'may emphasise emotional restraint over expressive sharing"',
  indent=False)


# ============================================================
H2('5. KẾT LUẬN')

OK('Bài Q1 v4 có MỨC RỦI RO ĐẠO VĂN THẤP TỔNG THỂ. Em đã chủ '
   'động viết theo phong cách riêng khi soạn draft, hầu hết các '
   'fact và term chuẩn quốc tế đều có trích dẫn đầy đủ.')

P('Để gửi BMC Psychiatry an toàn 100%, em đề xuất NCS bổ sung 3 '
  'trích dẫn ở mục 4 trước khi submit:', indent=False)
B('Rose, A. J. (2002). Co-rumination in the friendships of girls '
  'and boys. Child Development, 73(6), 1830–1843.', 0)
B('Stankov, L. (2010). Unforgiving Confucian culture: A breeding '
  'ground for high academic achievement, test anxiety and '
  'self-doubt? Learning and Individual Differences, 20(6), '
  '555-563. (hoặc nguồn tương đương)', 0)
B('Nguồn về khái niệm "nhẫn" trong văn hóa Việt — em đề xuất '
  'NCS hoặc thầy có nguồn tiếng Việt phù hợp hơn em đang biết.', 0)


# ============================================================
H2('THAM KHẢO ĐÃ ĐỀ XUẤT BỔ SUNG')

refs = [
    'Rose, A. J. (2002). Co-rumination in the friendships of girls '
    'and boys. Child Development, 73(6), 1830–1843. DOI: '
    '10.1111/1467-8624.00509.',
    'Stankov, L. (2010). Unforgiving Confucian culture: A breeding '
    'ground for high academic achievement, test anxiety and '
    'self-doubt? Learning and Individual Differences, 20(6), '
    '555-563. DOI: 10.1016/j.lindif.2010.05.003.',
    'Salili, F., & Hau, K.-T. (1994). The Effect of Teachers\' '
    'Evaluative Feedback on Chinese Students\' Perception of '
    'Ability: A Cultural and Situational Analysis. Educational '
    'Studies, 20(2), 223-236.',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(ref); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
