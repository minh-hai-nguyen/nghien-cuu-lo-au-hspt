# -*- coding: utf-8 -*-
"""Sinh bao cao ra soat sau Q1 paper outline.
- Logic flow review
- Every claim verification
- Issues + recommendations
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'RaSoat_Q1_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.3


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(15); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=False, color=None):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.first_line_indent = Cm(0.75)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic
    if color: r.font.color.rgb = color

def V(text):
    """Verified green"""
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('✓ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)

def X(text):
    """Issue red"""
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(2)
    r = p.add_run('⚠ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


# ============================================================
H1('BÁO CÁO RÀ SOÁT BÀI Q1 — LOGIC FLOW + CLAIMS')
P('Rà soát chi tiết: luồng dẫn dắt + mọi câu khẳng định + mọi dữ liệu',
  italic=True)
P('Ngày: 01/06/2026 — Phương án A (BMC Psychiatry)', italic=True)


# ============================================================
H1('PHẦN 1 — LOGIC FLOW REVIEW (luồng dẫn dắt)')

H2('1.1 Tổng quan dòng chảy 4 paragraphs Introduction')

H3('Hiện tại (outline v2)')
P('§1 → Burden globally + Asia + Vietnam gap')
P('§2 → Risk-protective framework + Vietnam methodology gap')
P('§3 → Gender patterns hypothesis')
P('§4 → Present study + 3 hypotheses')

H3('Vấn đề phát hiện')
X('Issue 1: §1 → §2 jump thiếu cầu nối. §1 nói burden + gap, §2 nhảy thẳng '
  'sang theoretical framework. Chưa giải thích WHY framework này phù hợp với gap.')
X('Issue 2: Mixed-methods rationale ABSENT trong toàn bộ Introduction. '
  'Contribution 3 nói mixed-methods, nhưng §1-§4 không justify why qualitative cần.')
X('Issue 3: Hypotheses chưa được phát biểu CỤ THỂ. Outline chỉ summary "H1, '
  'H2, H3" — cần spell out testable statements.')
X('Issue 4: Vietnamese cultural specificity (Confucian, family hierarchy) chưa '
  'được introduce trong §1-§4 — sẽ "đột ngột" xuất hiện ở Discussion.')

H3('Đề xuất sửa')
V('§1 (revised): Burden globally → Asia → Vietnam specifically; mention '
  'cultural specificity (Confucian academic pressure, family hierarchy)')
V('§2 (revised): WHY integrated R-P framework needed → vì previous studies '
  'tách Risk OR Protective riêng → không capture interaction; cite Lazarus & '
  'Folkman 1984 + Compas 2017 + Bronfenbrenner ecological')
V('§3 (revised): Gender + cultural mediation; cite Saikia 2023 explicit '
  '(30,0% boys > 18,9% girls, p=0,049 — Indian Northeast paradox); Wen 2020 '
  '(China rural similar paradox); McLean 2011 (Western evidence opposite)')
V('§4 (revised): Mixed-methods rationale + 3 explicit hypotheses:', )
V('   H1 (testable): "Three risk factors will show significant positive '
  'associations with all three anxiety disorder subtypes (GAD, SAD, SocAD)."',)
V('   H2 (testable): "Four protective factors will show significant negative '
  'associations with all three anxiety disorder subtypes."',)
V('   H3 (testable, novel): "Separation anxiety will demonstrate gender '
  'invariance, whereas GAD and SocAD will show gender differences (female > male)."',)

H2('1.2 Logic flow Methods → Results → Discussion')

H3('Hiện tại')
P('Methods chia 5 sections. Results 6 sections. Discussion 7 sections.')

H3('Vấn đề phát hiện')
X('Issue 5: Methods 2.4 chỉ list techniques (CFA, SEM, multi-group) mà KHÔNG '
  'specify fit indices threshold + invariance criteria. Reviewers BMC Psychiatry '
  'sẽ flag.')
X('Issue 6: Methods 2.3 (Qualitative interviews) — số lượng pp + framework analysis '
  'chưa rõ. Tạm flag NCS confirm sau, nhưng cần thêm: pp recruitment, sampling '
  'strategy, transcription, intercoder reliability.')
X('Issue 7: Methods thiếu chuyên mục "Mixed-methods integration design" — '
  'cần specify Convergent Parallel Design (Creswell & Plano Clark 2018).')
X('Issue 8: Results 3.4 R²=0,284 (risk) + R²=0,209 (protective) — đây là từ '
  '2 SEPARATE MODELS trong LA (Bảng 27 risk only, Bảng 30 protective only). '
  'INTEGRATED SEM (claim contribution 1) sẽ cần R² TỔNG HỢP với 7 predictors '
  'đồng thời. Hiện chưa có số này.')
X('Issue 9: Results 3.4 thiếu SUBTYPE-SPECIFIC β coefficients. LA có (ALHT→GAD '
  'β=0,510; ALHT→SocAD β=0,490; ALHT→SAD β=...) nhưng outline chỉ show TOTAL.')
X('Issue 10: Discussion 4.4 claim "cultural collectivism" + "developmental task '
  'universal" — KHÔNG có citations. Cần Triandis 1995 (collectivism) + Allen et al. '
  '2013 (development) hoặc tương đương.')

H3('Đề xuất sửa Methods')
V('2.4 Analytic strategy: thêm "Fit indices threshold (CFI ≥ 0,90; TLI ≥ 0,90; '
  'RMSEA ≤ 0,08; SRMR ≤ 0,08; Hu & Bentler 1999)"')
V('2.4: thêm "Multi-group invariance criteria: ΔCFI ≤ 0,01 (Cheung & Rensvold 2002)"')
V('2.5: thêm "Mixed-methods integration: Convergent Parallel Design '
  '(Creswell & Plano Clark 2018); joint display matrix"')

H3('Đề xuất sửa Results')
V('3.4: clarify "Integrated SEM model với 7 latent predictors → 3 anxiety '
  'subtypes (21 paths). Reporting CFI/TLI/RMSEA + β + 95% CI."')
V('3.4: thêm subtype-specific β table — verify từng giá trị trong LA Bảng 27-42')
V('3.4: R² breakdown per outcome (R² GAD, R² SAD, R² SocAD) — cần re-extract '
  'từ LA, KHÔNG dùng R²=0,284 + 0,209 cho integrated model')

H3('Đề xuất sửa Discussion')
V('4.4 (separation gender-invariance interpretation):', )
V('   "Cultural collectivism context (Triandis, 1995; Markus & Kitayama, 1991) — '
  'Vietnamese family hierarchy creates uniform separation experience across '
  'genders"',)
V('   "Developmental task pre-pubertal universality (Allen et al., 2013) — '
  'separation anxiety primarily childhood-onset, predates gender-differentiated '
  'social pressures emerging post-puberty"',)


d.add_page_break()


# ============================================================
H1('PHẦN 2 — CLAIMS VERIFICATION (từng câu khẳng định)')

H2('2.1 Introduction §1 claims')

V('Claim: "Xu 2021 (N=373.216 Chinese)" → VERIFIED. PDF: '
  'QT010_Xu_2021_China_LargestEpi.pdf, Abstract confirms 373,216 participants.')
V('Claim: "Anderson 2025 (narrative review)" → VERIFIED. PDF: '
  'QT014_Anderson_2025_Wiley_Narrative.pdf, J Child Adolesc Psychiatr Nurs.')
X('Claim: "GBD ASEAN 2025" → CHƯA RÕ. Có cite trong LA nhưng PDF source? '
  'Em đã thấy DICH_GBD_ASEAN_2025.docx + DICH_GBD_2021_ASEAN_Lancet_2025.md. '
  'CẦN VERIFY: là Lancet Regional Health 2025 hay GBD 2021 published 2025?')

H2('2.2 Introduction §2 claims')

V('Claim: "Lazarus & Folkman 1984" → standard theoretical reference. OK.')
V('Claim: "Compas 2017 (meta N=80,850)" → VERIFIED. PDF: '
  'Compas_2017_Coping_MetaAnalysis.pdf — Psychological Bulletin.')
V('Claim: "V-NAMHS 2022 (2,3% prevalence)" → VERIFIED. PDF: V-NAMHS_2022.pdf')
P('  Lưu ý: V-NAMHS dùng DISC-5, ngưỡng cắt khác DASS-21 → giải thích chênh '
  'lệch so với các nghiên cứu khác trong VN.', italic=True)

H2('2.3 Introduction §3 claims (gender)')

V('Claim: "McLean 2011 — female > male anxiety standard" → standard reference. OK.')
V('Claim: "Wen 2020 — China rural similar paradox (male > female)" → VERIFIED. '
  'PDF: QT008_Wen_2020_China_Rural_LPA.pdf (rural LMICs)')
V('Claim: "Saikia 2023 — men > women paradox (30,0% vs 18,9%, p=0,049)" → '
  'VERIFIED. PDF: 11_Saikia_2023_IJCM.pdf gốc page 5 — '
  '"Boy 54 (30.0%) vs Girl 34 (18.9%), p=0.049"')

H2('2.4 Methods claims')

V('Claim: "1.352 HS THCS Hà Nội" → VERIFIED LA + Thực trạng nguồn')
V('Claim: "614 nam / 738 nữ" → VERIFIED Thực trạng Bảng 4')
V('Claim: "8 validated scales" → VERIFIED LA P957-962')
V('Claim: "RCADS Chorpita 2000 7+4+4 items" → VERIFIED LA P959 — '
  '"từ 21 mục ban đầu, giữ 15 mục (lo âu lan tỏa: 7, lo âu chia ly: 4, '
  'lo âu xã hội: 4)"')
V('Claim: "ESSA Sun 2011 4 items" → VERIFIED LA P960 — "rút gọn còn 4 mục"')
V('Claim: "SAS-SV Kwon 2013 5 items" → VERIFIED LA P960 — "từ 10 mục giữ lại 5"')
V('Claim: "OBVQ Olweus 1996 8 items (4+4)" → VERIFIED LA P960 — '
  '"giữ lại 8 mục (mỗi dạng 4 mục)"')
V('Claim: "PSSM Goodenow 1993 7 items" → VERIFIED LA P961 — "từ 18 giữ 7"')
V('Claim: "MSPSS Zimet 1988 8 items (parental+peer subscales)" → VERIFIED '
  'LA P961 — "8 mục (mỗi nguồn 4 mục)"')
V('Claim: "RSES Rosenberg 1965 5 items" → VERIFIED LA P961 — "lựa chọn 5 mục"')
V('Claim: "Brief COPE Carver 1997 15 items (3 factors)" → VERIFIED LA P962')
V('Claim: "AMOS 31.0" → VERIFIED LA P247 — "AMOS 31.0"')

H2('2.5 Results claims (số liệu)')

V('Sample: 1.352 HS, 614 nam, 738 nữ, khối 6: 368, 7: 316, 8: 340, 9: 328 → '
  'TẤT CẢ VERIFIED Thực trạng Bảng 4')
V('GAD M=55,82, SAD M=25,06, SocAD M=48,41 → VERIFIED Thực trạng Bảng 1+2+3')
V('Gender F values: F=44,484 GAD; F=45,984 SocAD; F=29,642 Total; F=0,246 '
  'SAD → VERIFIED Thực trạng Bảng 4')
V('All p values: p<0,001 (3 đầu) + p=0,620 (SAD) → VERIFIED')
V('Grade trajectory SAD: 32,13 → 27,14 → 20,88 → 19,46 → VERIFIED '
  'Thực trạng Bảng 4 R4-R7')
V('β values: 0,533; 0,400; 0,276; -0,457; -0,273; -0,108; -0,015; +0,376 → '
  'TẤT CẢ VERIFIED LA β catalog')

X('Issue R²: 0,284 + 0,209 → đây là từ 2 MODELS RIÊNG trong LA. Integrated SEM '
  'sẽ cần R² mới. Cần action: re-run integrated SEM với data gốc hoặc '
  'modify claim thành "we report separate-model R² for comparability with LA"')

H2('2.6 Discussion claims')

V('Claim: "Pascoe 2020 SR — academic pressure → mental health" → VERIFIED. '
  'PDF: Pascoe_2020.pdf')
V('Claim: "Steare 2023 SR — 48/52 studies confirm positive association" → '
  'VERIFIED. PDF: Steare_2023_AcademicPressure_SR.pdf (Discussion: '
  '"Of the 52 studies included, 48 found evidence of a positive association")')
V('Claim: "Self-esteem ~85-89% magnitude of academic pressure" → VERIFIED '
  'LA P1416')
V('Claim: "Khung CT 8 nội dung" → VERIFIED LA P1397 + P1419')

X('Claim: "Cultural collectivism Vietnamese; developmental task universal" '
  '→ KHÔNG có citation hiện tại. Cần add: Triandis 1995 (collectivism), '
  'Markus & Kitayama 1991 (cultural self), Allen et al. 2013 (developmental task)')


d.add_page_break()


# ============================================================
H1('PHẦN 3 — TÓM TẮT 10 ISSUES + RECOMMENDATIONS')

t = d.add_table(rows=1, cols=3)
t.style = 'Light Grid Accent 1'
hdr = t.rows[0].cells
hdr[0].text = '#'; hdr[1].text = 'Issue'; hdr[2].text = 'Recommendation'
for c in hdr:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

issues = [
    ('1', 'Intro §1→§2 thiếu cầu nối logic',
     'Revise §1 với cultural specificity Vietnam; §2 mở đầu với "Therefore, an integrated framework is needed because..."'),
    ('2', 'Mixed-methods rationale absent Introduction',
     'Add 2 câu cuối §4: "Quantitative SEM alone cannot capture cultural/contextual mechanisms. Qualitative interviews illuminate β coefficients."'),
    ('3', 'Hypotheses không có phát biểu testable',
     'Spell out 3 hypotheses dạng "We hypothesized that..."'),
    ('4', 'Vietnamese cultural specificity vắng mặt Intro',
     'Add §1 closing line: "Vietnamese adolescents face unique Confucian academic pressures + family hierarchy — yet under-studied."'),
    ('5', 'Methods 2.4 thiếu fit indices threshold',
     'Add: "CFI ≥ 0,90; TLI ≥ 0,90; RMSEA ≤ 0,08; SRMR ≤ 0,08 (Hu & Bentler 1999); ΔCFI ≤ 0,01 for invariance (Cheung & Rensvold 2002)"'),
    ('6', 'Qualitative interviews chưa rõ',
     'NCS confirm số pp + sampling + intercoder reliability (Cohen κ); add to 2.3'),
    ('7', 'Methods thiếu MM integration design',
     'Add: "Convergent Parallel Design (Creswell & Plano Clark 2018) với joint display matrix"'),
    ('8', 'R²=0,284 + 0,209 KHÔNG phải integrated R²',
     'Option A: re-run integrated SEM với 7 predictors → 1 R²; Option B: clarify "separate model R² presented for comparability with parent dissertation"'),
    ('9', 'Subtype-specific β missing',
     'Add full β table 7×3 (7 predictors × 3 outcomes). Verify từng cell trong LA Bảng 27-42'),
    ('10', 'Discussion 4.4 thiếu cultural/developmental citations',
     'Add: Triandis 1995; Markus & Kitayama 1991; Allen et al. 2013'),
]
for issue in issues:
    row = t.add_row().cells
    for i, txt in enumerate(issue):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)


d.add_page_break()


# ============================================================
H1('PHẦN 4 — VERIFICATION SUMMARY TABLE')

t2 = d.add_table(rows=1, cols=3)
t2.style = 'Light Grid Accent 1'
hdr2 = t2.rows[0].cells
hdr2[0].text = 'Category'
hdr2[1].text = 'Verified'
hdr2[2].text = 'Issues'
for c in hdr2:
    for p in c.paragraphs:
        for r in p.runs:
            r.font.bold = True; r.font.size = Pt(11)

ver_summary = [
    ('Introduction citations', '8/9 verified', '1 (GBD ASEAN 2025 — cần clarify nguồn)'),
    ('Methods sample data', '5/5 verified', '0'),
    ('Methods 8 scales', '8/8 verified', '0'),
    ('Results descriptive M/SD', '6/6 verified', '0'),
    ('Results F + p values', '8/8 verified', '0'),
    ('Results SEM β coefficients', '7/7 verified', '0'),
    ('Results R²', 'PARTIAL', '1 (R² is from separate models, not integrated)'),
    ('Discussion claims', '4/6 verified', '2 (Cultural collectivism + Developmental task — missing citations)'),
    ('17 References PDFs', '17/17 PDF exists', '0 (after Saikia found)'),
    ('Logic flow Introduction', '—', '4 issues (jump, MM rationale, hypotheses, culture)'),
    ('Logic flow Methods', '—', '3 issues (fit indices, qualitative detail, MM design)'),
    ('Logic flow Results', '—', '2 issues (R² interpretation, subtype β)'),
    ('Logic flow Discussion', '—', '1 issue (cultural/developmental citations)'),
]
for row_data in ver_summary:
    row = t2.add_row().cells
    for i, txt in enumerate(row_data):
        row[i].text = txt
        for p in row[i].paragraphs:
            for r in p.runs:
                r.font.size = Pt(10)
                if 'verified' in txt.lower() and 'PARTIAL' not in txt:
                    r.font.color.rgb = RGBColor(0x00, 0x70, 0x00)
                elif 'issues' in txt.lower() or 'PARTIAL' in txt:
                    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)


H2('Tổng kết em đề xuất')

P('Q1 outline v2 có CONTENT VERIFICATION rất tốt (most claims có evidence). '
  'Nhưng LOGIC FLOW + METHODOLOGICAL COMPLETENESS còn 10 issues cần fix '
  'trước khi viết draft.', italic=True)

P('Critical priority (BLOCKING):', )
P('  • Issue 8 (R² integrated vs separate) — cần thầy + NCS quyết: re-run '
  'integrated SEM hay clarify limitation', indent=False)
P('  • Issue 6 (Qualitative confirm) — NCS confirm số pp + sampling', indent=False)
P('  • Issue 9 (Subtype-specific β table) — em re-extract từ LA Bảng 27-42', indent=False)

P('High priority (can fix in v3 outline):', )
P('  • Issue 1, 2, 3, 4 (Introduction logic flow + hypotheses)', indent=False)
P('  • Issue 5, 7 (Methods fit indices + MM design)', indent=False)
P('  • Issue 10 (Discussion citations cultural)', indent=False)

P('Em đề xuất build outline v3 sau khi thầy + NCS quyết Issue 6 + 8. '
  'V3 sẽ fix tất cả 10 issues + viết đầy đủ hypotheses + fit indices threshold.',
  italic=True)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
