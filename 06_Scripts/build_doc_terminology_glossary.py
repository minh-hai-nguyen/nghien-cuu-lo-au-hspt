# -*- coding: utf-8 -*-
"""
Generate comprehensive glossary: tất cả dữ liệu/tham số/thang đo/phương pháp trong 68 papers.
Output: 01_Bao-cao/Giai_thich_toan_bo_thuat_ngu_CSDL.docx
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '01_Bao-cao', 'Giai_thich_toan_bo_thuat_ngu_CSDL.docx')

doc = Document()
doc.styles['Normal'].font.name = 'Times New Roman'
doc.styles['Normal'].font.size = Pt(11)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def PRun(parts):
    p = doc.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text)
        r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 11))
        if opts.get('bold'): r.bold = True
        if opts.get('italic'): r.italic = True
        if opts.get('red'): r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
        elif opts.get('color'): r.font.color.rgb = opts['color']
    return p

def H1(t): return P(t, bold=True, size=16, color=RGBColor(0x1F, 0x3A, 0x68), align=WD_ALIGN_PARAGRAPH.CENTER)
def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))
def H4(t): return P(t, bold=True, italic=True, size=11, color=RGBColor(0x2E, 0x54, 0x8B))

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(9.5)
    return t

def note_box(text, color_hex='FFF3E0'):
    t = doc.add_table(rows=1, cols=1)
    t.style = 'Table Grid'
    c = t.rows[0].cells[0]; c.text = ''
    r = c.paragraphs[0].add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(10.5); r.italic = True
    set_cell_bg(c, color_hex)

# ============================================================
# TRANG BÌA
# ============================================================
H1('GIẢI THÍCH TOÀN BỘ THUẬT NGỮ')
H1('TRONG CƠ SỞ DỮ LIỆU DỰ ÁN LO ÂU')
P('Từ điển tra cứu mọi thang đo / tham số / phương pháp / thiết kế NC xuất hiện trong 68 bài canonical (VN001–VN030 + QT001–QT051)',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=11, color=RGBColor(0x66, 0x66, 0x66))
P('Nhóm dự án Lo âu | Tháng 04/2026',
  italic=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=10, color=RGBColor(0x66, 0x66, 0x66))
P()

note_box(
    'CÁCH DÙNG: Đây là tài liệu tra cứu (reference). Khi đọc báo cáo mà gặp thuật ngữ lạ '
    '(vd "SCARED", "RoB 2", "τ²", "SUCRA") → mở doc này, Ctrl+F tìm, đọc giải nghĩa ngắn + '
    'papers áp dụng. Cấu trúc 7 phần: A (Thang đo), B (Chẩn đoán), C (Thiết kế), D (Can thiệp), '
    'E (Tham số thống kê), F (Khái niệm lâm sàng), G (Chuẩn báo cáo).',
    color_hex='FFFDE7')
P()

# ============================================================
# MỤC LỤC
# ============================================================
H2('MỤC LỤC')
P('A. THANG ĐO / INSTRUMENTS (6 tiểu mục)')
P('B. CHẨN ĐOÁN TÂM THẦN (DSM-5/ICD-11)')
P('C. THIẾT KẾ NGHIÊN CỨU (5 loại chính)')
P('D. PHƯƠNG PHÁP CAN THIỆP (7 nhóm)')
P('E. THAM SỐ THỐNG KÊ (6 nhóm)')
P('F. KHÁI NIỆM LÂM SÀNG (6 nhóm)')
P('G. CHUẨN BÁO CÁO KHOA HỌC (5 chuẩn chính)')
P('H. BẢNG ĐỐI CHIẾU NHANH: Papers × Instruments × Designs')
doc.add_page_break()

# ============================================================
# PHẦN A — THANG ĐO
# ============================================================
H1('PHẦN A. THANG ĐO / INSTRUMENTS')
P('Các công cụ đo lường xuất hiện trong thư viện dự án.', italic=True)
P()

# A.1 Thang đo LO ÂU
H2('A.1. Thang đo LO ÂU (anxiety measures)')
MakeTable(
    ['Thang đo', 'Tên đầy đủ', 'Đặc điểm', 'Papers dùng'],
    [
        ('GAD-7',
         'Generalized Anxiety Disorder 7-item Scale (Spitzer et al. 2006)',
         '7 items, 0–21 điểm. Cut-off: 5 (nhẹ), 10 (vừa), 15 (nặng). Sàng lọc lo âu tổng quát ở người lớn và VTN. Đã validate tiếng Việt.',
         'VN001 (Hoa 2024), V-NAMHS (đo phụ huynh)'),
        ('SCARED',
         'Screen for Child Anxiety Related Emotional Disorders',
         '41 items (bản đầy đủ), 5 subscale: panic, GAD, separation anxiety, social anxiety, school avoidance. Cho trẻ 8–18 tuổi. Tự báo cáo + parent-report.',
         'Nhiều NC quốc tế trong QT029, QT039 NMA'),
        ('STAI-C',
         'State-Trait Anxiety Inventory for Children',
         '40 items, đo state (trạng thái hiện tại) + trait (đặc điểm nhân cách). Cho trẻ 9–12 tuổi. Spielberger 1973.',
         'Cộng sự VN (Nguyễn Thị Vân 2020 STAI TPHCM)'),
        ('HAM-A',
         'Hamilton Anxiety Rating Scale',
         '14 items, clinician-rated (không phải self-report). 0–56 điểm. Chuẩn vàng cho đánh giá lâm sàng lo âu. Đặc biệt phổ biến trong NC điều trị người lớn.',
         'Trần Nguyễn Ngọc 2018 (luận án Bạch Mai)'),
        ('LSAS-SR',
         'Liebowitz Social Anxiety Scale — Self Report',
         '24 items, đo social anxiety specifically. Đo 2 dimension: fear + avoidance. Cut-off: ≥ 30 = social anxiety lâm sàng.',
         'QT045 (Matsumoto Japan iCBT) — primary outcome'),
        ('BAI',
         'Beck Anxiety Inventory',
         '21 items, 0–63. Tập trung triệu chứng cơ thể của lo âu (đau ngực, run, sợ mất kiểm soát). Cut-off: 8 (nhẹ), 16 (vừa), 26 (nặng).',
         'Một số NC quốc tế trong QT028 review Zugman'),
    ])
P()

# A.2 Thang đo TRẦM CẢM
H2('A.2. Thang đo TRẦM CẢM (depression measures)')
MakeTable(
    ['Thang đo', 'Tên đầy đủ', 'Đặc điểm', 'Papers dùng'],
    [
        ('PHQ-9',
         'Patient Health Questionnaire 9-item (Kroenke et al. 2001)',
         '9 items, 0–27. Cut-off: 5 (nhẹ), 10 (vừa), 15 (vừa-nặng), 20 (nặng). Sàng lọc trầm cảm tổng quát. Đã validate tiếng Việt. ĐI KÈM GAD-7 nhưng KHÔNG gộp chung.',
         'V-NAMHS (đo phụ huynh); nhiều NC VN'),
        ('CESD-R',
         'Center for Epidemiologic Studies Depression Scale — Revised',
         '20 items, 0–80. Cut-off ≥ 16 = trầm cảm lâm sàng. Eaton et al. 2004 revision. Đo trong tuần qua.',
         'VN030 Happy House (primary outcome)'),
        ('BDI / BDI-II',
         'Beck Depression Inventory (II = 2nd edition, Beck et al. 1996)',
         '21 items, 0–63. Cut-off: 14 (nhẹ), 20 (vừa), 29 (nặng). Phổ biến nhất trong NC tâm lý học lâm sàng.',
         'Một số NC trong QT028, QT029 reviews'),
        ('PHQ-A',
         'Patient Health Questionnaire — Adolescent (11–17 tuổi)',
         'Bản adapt của PHQ-9 cho VTN. 9 items.',
         'Một số NC quốc tế (cite trong review)'),
        ('CES-DC',
         'CES-D Scale for Children',
         '20 items, cho trẻ 6–17. Weissman 1980. 0–60 điểm.',
         'Một số NC VN cộng sự'),
    ])
P()

# A.3 Thang đo ĐA CHIỀU
H2('A.3. Thang đo ĐA CHIỀU (multi-dimensional)')
MakeTable(
    ['Thang đo', 'Tên đầy đủ', 'Đặc điểm', 'Papers dùng'],
    [
        ('DASS-21',
         'Depression Anxiety Stress Scale — 21 items (Lovibond & Lovibond 1995)',
         '21 items = 3 subscale × 7 items: trầm cảm, lo âu, stress. Mỗi subscale 0–42 (× 2 với DASS-21 để so với DASS-42).',
         'VN029 Duong TPHCM (lo âu 50,3 %); nhiều NC VN'),
        ('DASS-Y',
         'DASS-21 phiên bản Youth (adapted cho VTN)',
         'Tương tự DASS-21 nhưng diễn đạt phù hợp tuổi 12–18. Cùng 3 subscale.',
         'VN020 Trần Hồ Vĩnh Lộc TPHCM'),
        ('DASS-42',
         'DASS phiên bản đầy đủ — 42 items',
         'Tiền thân của DASS-21. 42 items = 3 subscale × 14.',
         'VN (Trần Thị Mỵ Lương 2020 DASS-42 THPT Chuyên)'),
        ('SDQ-25',
         'Strengths and Difficulties Questionnaire — 25 items (Goodman 1997)',
         '5 subscale: emotional, conduct, hyperactivity, peer, prosocial. Self-report (11–17) + parent/teacher (4–17). Cut-off tổng: ≥ 17 = cao nguy cơ. VALIDATE tiếng Việt bởi Weiss 2014.',
         'VN022 UNICEF (26,1 % nguy cơ cao); VN002 VNAMHS (tham chiếu)'),
        ('PSC-17',
         'Pediatric Symptom Checklist — 17 items',
         '17 items, parent-report, 4–17 tuổi. 3 subscale: internalising, externalising, attention. Cut-off: 15 = rủi ro.',
         'V-NAMHS (đo từ phụ huynh)'),
        ('CBCL',
         'Child Behaviour Checklist (Achenbach)',
         '118 items, parent-report. Internalising + Externalising + 8 DSM-oriented scales. T-score ≥ 60 = borderline; ≥ 64 = clinical.',
         'Weiss 2014 (cite trong VN002)'),
        ('YSR',
         'Youth Self-Report (Achenbach)',
         '112 items, self-report 11–18 tuổi. Parallel với CBCL.',
         'Weiss 2014 (cite trong VN002)'),
    ])
P()

# A.4 Công cụ CHẨN ĐOÁN
H2('A.4. Công cụ CHẨN ĐOÁN (diagnostic instruments)')
MakeTable(
    ['Công cụ', 'Tên đầy đủ', 'Đặc điểm', 'Papers dùng'],
    [
        ('DISC-5',
         'Diagnostic Interview Schedule for Children, Version 5 (Bitsko 2019; Shaffer 2000)',
         'Phỏng vấn cấu trúc hoàn toàn cho trẻ 6–18. Lay-interviewer có thể dùng. Module theo DSM-5. Cho 20+ rối loạn. KHÁC thang sàng lọc ở chỗ cho CHẨN ĐOÁN DSM-5 chính thức.',
         'VN002 V-NAMHS (ĐẦU TIÊN tại VN) — 6 module: social phobia, GAD, MDD, PTSD, conduct, ADHD'),
        ('K-SADS',
         'Kiddie Schedule for Affective Disorders and Schizophrenia',
         'Phỏng vấn chẩn đoán semi-structured cho trẻ 6–18. Clinician-administered. Cho mọi rối loạn Axis I DSM.',
         'Một số NC quốc tế'),
        ('MINI / MINI-KID',
         'Mini International Neuropsychiatric Interview (Sheehan 1998)',
         'Phỏng vấn chẩn đoán ngắn (~15 phút) cho người lớn / trẻ. Cho DSM-IV/5 + ICD-10. Rất phổ biến ở LMIC.',
         'Một số NC VN + khu vực'),
        ('CIDI',
         'Composite International Diagnostic Interview (WHO)',
         'Fully structured, lay interviewer. Chuẩn WHO World Mental Health Surveys. Cho DSM-5/ICD-11. Dài hơn DISC-5.',
         'Erskine 2021 (NAMHS protocol)'),
    ])
P()

# A.5 Thang đo TRƯỜNG HỌC + LIÊN QUAN
H2('A.5. Thang đo TRƯỜNG HỌC + MÔI TRƯỜNG')
MakeTable(
    ['Thang đo', 'Tên đầy đủ', 'Đặc điểm', 'Papers dùng'],
    [
        ('MDS3',
         'Maryland Safe and Supportive Schools Climate Survey',
         '56 items cho HS. 3 domain: safety, engagement, environment. Adapt bởi Bradshaw 2014.',
         'VN022 UNICEF (khí hậu trường)'),
        ('ESSA',
         'Engagement with School Scale for Adolescents (Sun et al. 2011)',
         'Đo school engagement — behavioural, emotional, cognitive.',
         'VN022 UNICEF'),
        ('RCBI',
         'Revised Cyber Bullying Inventory (Topcu & Erdur-Baker)',
         'Đo tần suất + distress của cyberbullying. Validate VN bởi Cong et al. 2018.',
         'VN022 UNICEF (cyberbullying 52,2 % bị)'),
        ('ACE / ACEs',
         'Adverse Childhood Experiences questionnaire (Felitti 1998)',
         '10–17 items. Đo các trải nghiệm bất lợi thời thơ ấu: lạm dụng, bỏ rơi, rối loạn chức năng gia đình.',
         'VN002 V-NAMHS'),
        ('GEAS Family Connectedness',
         'Global Early Adolescent Study — Family Connectedness module',
         'Đo mức độ kết nối VTN với cha mẹ/người chăm sóc.',
         'VN002 V-NAMHS'),
        ('Bullying frequency (general)',
         'Various instruments (Olweus Bully/Victim Questionnaire etc.)',
         'Đo tần suất bị bắt nạt / bắt nạt người khác. Nhiều dạng: victim, perpetrator, bystander.',
         'VN022, VN029, nhiều NC'),
    ])
P()

# A.6 Tự trọng, Resilience
H2('A.6. Thang đo TỰ TRỌNG, RESILIENCE, COPING')
MakeTable(
    ['Thang đo', 'Tên đầy đủ', 'Đặc điểm', 'Papers dùng'],
    [
        ('Rosenberg Self-Esteem',
         'Rosenberg Self-Esteem Scale (1965)',
         '10 items, 0–30. Đo tự trọng tổng quát. Chuẩn nhất trong NC.',
         'VN002 V-NAMHS'),
        ('CD-RISC',
         'Connor-Davidson Resilience Scale',
         '25 items (+ versions 10, 2 items). Đo khả năng phục hồi tâm lý.',
         'Một số NC QT, cộng sự VN'),
        ('CSES',
         'Coping Self-Efficacy Scale',
         '26 items. 3 subscale: problem-focused, stop unpleasant emotions, social support.',
         'VN030 Happy House (secondary outcome)'),
        ('MHC-SF',
         'Mental Health Continuum — Short Form (Keyes)',
         '14 items, 3 subscale: emotional, social, psychological wellbeing.',
         'VN030 Happy House'),
        ('School Connectedness',
         'Various instruments',
         'Đo mức độ kết nối của HS với trường học.',
         'VN030, VN027'),
    ])
P()
doc.add_page_break()

# ============================================================
# PHẦN B — CHẨN ĐOÁN
# ============================================================
H1('PHẦN B. CHẨN ĐOÁN TÂM THẦN (DSM-5 / ICD-11)')

H2('B.1. Rối loạn LO ÂU (Anxiety Disorders)')
MakeTable(
    ['Viết tắt', 'Tên đầy đủ', 'Đặc điểm chẩn đoán chính'],
    [
        ('GAD',
         'Generalized Anxiety Disorder (Rối loạn lo âu lan toả)',
         'Lo âu quá mức về nhiều sự kiện, ≥ 6 tháng, khó kiểm soát, kèm 3+ triệu chứng (căng thẳng cơ, mệt mỏi, khó tập trung, cáu gắt, rối loạn giấc ngủ). DSM-5 code F41.1.'),
        ('SAD',
         'Social Anxiety Disorder (Rối loạn lo âu xã hội)',
         'Sợ tình huống xã hội ≥ 6 tháng, sợ bị đánh giá tiêu cực, né tránh hoặc chịu đựng với distress. Ở VTN: phải trong peer-setting, không chỉ với người lớn. DSM-5 F40.10.'),
        ('Social Phobia',
         '(Tên cũ của SAD — dùng trong DISC-5)',
         'Tương đương SAD. DISC-5 vẫn dùng "social phobia".'),
        ('Panic Disorder',
         'Rối loạn hoảng sợ',
         'Cơn hoảng sợ đột ngột + ≥ 1 tháng lo về cơn tiếp theo. DSM-5 F41.0.'),
        ('Specific Phobia',
         'Ám ảnh sợ đặc hiệu',
         'Sợ đối tượng/tình huống cụ thể (nhện, máy bay, máu...). DSM-5 F40.xxx.'),
        ('Separation Anxiety',
         'Lo âu chia ly',
         'Phổ biến ở trẻ nhỏ — sợ xa người gắn bó chính. DSM-5 F93.0.'),
    ])
P()

H2('B.2. Rối loạn TRẦM CẢM (Depressive Disorders)')
MakeTable(
    ['Viết tắt', 'Tên đầy đủ', 'Đặc điểm'],
    [
        ('MDD',
         'Major Depressive Disorder (Rối loạn trầm cảm nặng)',
         '≥ 5 trong 9 triệu chứng (tâm trạng buồn, anhedonia, thay đổi cân nặng, rối loạn giấc ngủ, tâm động cơ thể, mệt, cảm giác vô dụng, khó tập trung, ý nghĩ tự sát) ≥ 2 tuần. DSM-5 F32.x.'),
        ('PDD',
         'Persistent Depressive Disorder (Dysthymia)',
         'Trầm cảm mạn ≥ 1 năm (trẻ em/VTN) hoặc ≥ 2 năm (người lớn). Triệu chứng nhẹ hơn MDD nhưng kéo dài. DSM-5 F34.1.'),
    ])
P()

H2('B.3. PTSD — Posttraumatic Stress Disorder')
P('DSM-5 F43.10. Sau sang chấn lớn → 4 nhóm triệu chứng: (1) Intrusion (xâm nhập ký ức); (2) '
  'Avoidance (né tránh); (3) Negative alterations in cognition/mood; (4) Arousal/reactivity. '
  '≥ 1 tháng. Có phiên bản PTSD cho trẻ < 6 tuổi.')
P()

H2('B.4. ADHD — Attention-Deficit/Hyperactivity Disorder')
P('DSM-5 F90.x. ≥ 6 triệu chứng inattention (khó tập trung, quên, mất đồ, dễ mất chú ý) hoặc '
  'hyperactivity/impulsivity (quẫy, chạy nhảy, nói nhiều, cắt lời), ≥ 6 tháng, ở ≥ 2 môi trường, '
  'khởi phát trước 12 tuổi. 3 presentation: inattentive, hyperactive-impulsive, combined.')
P()

H2('B.5. Rối loạn HÀNH VI (Conduct Disorder)')
P('DSM-5 F91.x. Mô hình hành vi lặp lại vi phạm quyền người khác hoặc chuẩn mực xã hội: '
  'aggression, destruction of property, deceitfulness/theft, serious rules violation. ≥ 3 '
  'trong 15 triệu chứng, ≥ 12 tháng.')
P()

H2('B.6. HỆ CHUẨN: DSM-5 vs ICD-11')
MakeTable(
    ['Chuẩn', 'Do ai', 'Phổ biến ở đâu', 'Điểm quan trọng'],
    [
        ('DSM-5',
         'American Psychiatric Association (APA, 2013)',
         'Mỹ + nghiên cứu quốc tế (tạp chí Mỹ/Anh)',
         'Chẩn đoán theo criteria list (≥ N triệu chứng). DISC-5 theo DSM-5.'),
        ('DSM-5-TR',
         'APA 2022 — Text Revision',
         'Cập nhật dựa trên DSM-5',
         'Thêm PGD (Prolonged Grief Disorder); làm rõ criteria.'),
        ('ICD-11',
         'World Health Organization (WHO, 2022)',
         'Châu Âu + WHO-endorsed countries (VN)',
         'Chẩn đoán theo prototype + severity. Khác DSM-5 ở một số category.'),
        ('ICD-10',
         'WHO — vẫn dùng ở nhiều quốc gia đang chuyển sang ICD-11',
         'VN hiện dùng chủ yếu ICD-10 (F00–F99 mental disorders)',
         'F41.1 = GAD; F32 = MDD; F43.1 = PTSD.'),
    ])
P()
doc.add_page_break()

# ============================================================
# PHẦN C — THIẾT KẾ NGHIÊN CỨU
# ============================================================
H1('PHẦN C. THIẾT KẾ NGHIÊN CỨU')

H2('C.1. Nghiên cứu CẮT NGANG (Cross-sectional)')
P('Thu thập dữ liệu tại 1 thời điểm. ', bold=True)
P('Đặc điểm: Rẻ, nhanh. CHỈ cho biết tương quan, KHÔNG biết nhân quả.',
  italic=True)
P('Ví dụ trong thư viện:', bold=True)
P('• VN001 (Hoa 2024) — GAD-7 trên 3.910 HS Hà Nội', size=10)
P('• VN002 V-NAMHS (2022) — 5.996 cặp phụ huynh-VTN, 38 tỉnh', size=10)
P('• VN022 UNICEF (2022) — 668 HS 4 tỉnh', size=10)
P('• VN029 Duong (2025) — 2.631 HS THPT TPHCM', size=10)
P()

H2('C.2. Nghiên cứu COHORT / LONGITUDINAL')
P('Theo dõi cùng một nhóm qua thời gian (prospective) hoặc nhìn ngược về quá khứ (retrospective). ',
  bold=True)
P('Đặc điểm: Xác lập trình tự thời gian, cho phép suy đoán nhân quả TỐT HƠN cắt ngang.',
  italic=True)
P('Ví dụ: QT021 Brunborg 2025 Norway — panel study social media', size=10)
P()

H2('C.3. RCT (Randomized Controlled Trial)')
P('Phân bổ NGẪU NHIÊN người tham gia vào nhóm can thiệp vs đối chứng. CHUẨN VÀNG cho nhân quả.',
  bold=True)

MakeTable(
    ['Loại RCT', 'Đặc điểm', 'Ví dụ trong thư viện'],
    [
        ('RCT cá nhân',
         'Cá nhân được ngẫu nhiên',
         'QT043 Bress 2024 Maya app (young adults)'),
        ('Cluster RCT',
         'Cụm (trường, lớp) được ngẫu nhiên, không phải cá nhân',
         'QT038 De Silva 2024 Sri Lanka (36 trường)'),
        ('Parallel arm',
         'Mỗi người trong 1 arm (CT hoặc ĐC)',
         'Đa số RCT trong thư viện'),
        ('Crossover',
         'Mỗi người trải nghiệm cả 2 arm với washout',
         'Không có trong thư viện (hiếm với can thiệp CBT)'),
        ('Two-arm',
         '2 nhóm so sánh',
         'VN030 Happy House (nhưng non-randomized)'),
        ('Multi-arm',
         '3+ nhóm so sánh',
         'Một số NC trong QT029 NMA'),
    ])
P()

H2('C.4. Quasi-experimental (Bán thực nghiệm)')
P('Có nhóm đối chứng NHƯNG KHÔNG phân bổ ngẫu nhiên. Ví dụ: trường tham gia vs không tham gia.',
  italic=True)
P('Ví dụ: VN030 Happy House Cluster CONTROLLED (không randomized do hành chính)', size=10)
P()

H2('C.5. SR, MA, NMA, Scoping Review')
P('Xem chi tiết trong document "Giải thích MA vs NMA". Tóm tắt:', italic=True)
MakeTable(
    ['Loại review', 'Đặc điểm', 'Output'],
    [
        ('SR (Systematic Review)',
         'Tổng quan hệ thống theo PRISMA',
         'Narrative synthesis — không pool effect sizes'),
        ('MA (Meta-Analysis)',
         'Pool effect sizes từ nhiều NC về 1 so sánh',
         '1 SMD + 95 % CI/CrI'),
        ('NMA (Network MA)',
         'So sánh nhiều can thiệp cùng lúc qua mạng',
         'Ma trận SMD + SUCRA xếp hạng'),
        ('Scoping Review',
         'Khám phá phạm vi bằng chứng (không evaluate efficacy)',
         'Bản đồ bằng chứng, gap identification'),
    ])
P()

H2('C.6. Case-control + Case report/series')
P('• Case-control: so nhóm có bệnh vs không có bệnh, hỏi ngược về exposure trước đó', size=10)
P('• Case report: 1 trường hợp lâm sàng đặc biệt', size=10)
P('• Case series: nhiều trường hợp tương tự', size=10)
P('(Hiếm trong thư viện chủ yếu focus intervention/prevalence)', italic=True, size=10)
P()
doc.add_page_break()

# ============================================================
# PHẦN D — CAN THIỆP
# ============================================================
H1('PHẦN D. PHƯƠNG PHÁP CAN THIỆP')

H2('D.1. Họ CBT (Cognitive Behavioral Therapy)')
MakeTable(
    ['Loại CBT', 'Đặc điểm', 'Papers'],
    [
        ('CBT truyền thống (individual / face-to-face)',
         'Cá nhân gặp trị liệu viên. 8–20 buổi. Gồm psychoeducation, cognitive restructuring, exposure, behavioural activation.',
         'Walkup 2008 CAMS, nhiều NC trong QT029 NMA'),
        ('gCBT (Group CBT)',
         'Nhóm 6–10 người + trị liệu viên. Cùng nội dung CBT nhưng delivered theo nhóm.',
         'QT039 Xian 2024 (SUCRA 68,4 %)'),
        ('iCBT (Internet CBT)',
         'Qua website/platform. Có 2 kiểu: GUIDED (có coach/therapist hỗ trợ) và UNGUIDED (tự học).',
         'QT045 Matsumoto, QT040 Walder, QT039 Xian (SUCRA 71,2 %)'),
        ('Mobile CBT',
         'Qua smartphone app. Thường là iCBT nhưng delivered qua mobile.',
         'QT043 Bress Maya, QT050 Qiaochu & Wang'),
        ('CA-CBT',
         'Culturally Adapted CBT — CBT điều chỉnh văn hoá (peripheral hoặc core adaptation).',
         'QT037 Praptomojati & Hartanto ĐNA'),
        ('DBT (Dialectical Behavior Therapy)',
         'Biến thể CBT phát triển bởi Linehan. Cho BPD, tự sát, cảm xúc không ổn định.',
         'Không có trong thư viện lõi nhưng cite trong review'),
        ('ACT (Acceptance and Commitment Therapy)',
         'Third-wave CBT. Dựa acceptance + mindfulness + commitment to values.',
         'Nhiều NC quốc tế trong QT029 NMA'),
    ])
P()

H2('D.2. Mindfulness-based interventions')
MakeTable(
    ['Chương trình', 'Đặc điểm', 'Papers'],
    [
        ('MBSR',
         'Mindfulness-Based Stress Reduction (Kabat-Zinn) — 8 tuần × 2,5 giờ',
         'Cite trong nhiều review'),
        ('MBCT',
         'Mindfulness-Based Cognitive Therapy — MBSR + CBT cho relapse prevention trầm cảm',
         'Cite trong meta-analyses'),
        ('MYRIAD',
         'Universal mindfulness program UK (Kuyken 2022 — 8.376 HS / 85 trường)',
         'Kuyken 2022 (không có PDF nhưng cite trong báo cáo 10/04)'),
        ('.b / Paws b',
         'Mindfulness in Schools Project (UK) — adapt cho HS',
         'Cite trong QT042 BESST editorial'),
    ])
P()

H2('D.3. Resilience programs')
MakeTable(
    ['Chương trình', 'Đặc điểm', 'Papers'],
    [
        ('RAP-A',
         'Resourceful Adolescent Program (Shochet) — resilience CBT-based, 11 buổi',
         'VN030 Happy House (adapt RAP-A văn hoá VN)'),
        ('FRIENDS for Life',
         'Anxiety prevention program cho trẻ em (Barrett) — 10 buổi',
         'Cite trong nhiều review quốc tế'),
        ('Penn Resiliency Program',
         'CBT-based resilience prevention (Seligman) — 12 buổi',
         'Cite trong QT029, QT049'),
        ('Cool Kids',
         'Anxiety program Australian (Rapee) — 10 buổi + parent component',
         'Cite trong QT028 Zugman review'),
    ])
P()

H2('D.4. Digital Mental Health Interventions (DMHI)')
P('Bao gồm iCBT, mobile CBT, chatbot, VR. Đặc điểm: rẻ, scalable, phù hợp LMIC.',
  italic=True)
P()
P('Papers: QT040 Walder DMHI SAD, QT043 Bress Maya, QT045 Matsumoto Japan iCBT, '
  'QT050 Qiaochu Mobile CBT.', size=10)
P()

H2('D.5. MHPSS (Mental Health and Psychosocial Support)')
P('Khung của WHO/IASC — gồm 4 tầng: (1) services; (2) community supports; (3) family supports; '
  '(4) basic services. Rộng hơn chỉ can thiệp lâm sàng.')
P('Papers: QT051 Menon LMIC scoping, VN022 UNICEF, VN030 Happy House', size=10)
P()

H2('D.6. Pharmacological (Thuốc)')
MakeTable(
    ['Nhóm thuốc', 'Loại', 'Papers'],
    [
        ('SSRI',
         'Selective Serotonin Reuptake Inhibitors (Fluoxetine, Sertraline, Escitalopram)',
         'Walkup 2008 CAMS (CBT+Sertraline); cite trong QT028 Zugman'),
        ('SNRI',
         'Serotonin-Norepinephrine RI (Venlafaxine, Duloxetine)',
         'Cite trong review'),
        ('Benzodiazepine',
         'Diazepam, Lorazepam — ngắn hạn, tránh cho VTN do addiction risk',
         'Hiếm trong NC hiện đại cho VTN'),
    ])
P()

H2('D.7. Non-pharmacological khác')
MakeTable(
    ['Phương pháp', 'Đặc điểm', 'Papers'],
    [
        ('Yoga',
         'Kết hợp tư thế, hơi thở, thiền. Evidence cho lo âu nhẹ–vừa.',
         'Trần Nguyễn Ngọc 2018 Bạch Mai'),
        ('Physical Exercise (PE)',
         'Tập thể dục ≥ 3×/tuần × 30 phút. SUCRA 0,51 cho anxiety children.',
         'QT029 Li 2025 NMA'),
        ('Relaxation Therapy',
         'Progressive muscle relaxation (Jacobson), breathing exercises',
         'Trần Nguyễn Ngọc 2018'),
        ('Biofeedback',
         'Phản hồi sinh học (HRV, GSR) để tự điều chỉnh',
         'Hiếm trong thư viện'),
        ('Art Therapy / Music Therapy',
         'Creative arts để expression + regulation',
         'Một số NC quốc tế không phải lõi'),
    ])
P()
doc.add_page_break()

# ============================================================
# PHẦN E — THAM SỐ THỐNG KÊ
# ============================================================
H1('PHẦN E. THAM SỐ THỐNG KÊ')

H2('E.1. Thống kê MÔ TẢ')
MakeTable(
    ['Ký hiệu', 'Ý nghĩa', 'Ví dụ'],
    [
        ('n / N',
         'Cỡ mẫu (sample size). n thường cho subgroup, N cho total.',
         'N = 5.996 cặp (V-NAMHS)'),
        ('M hoặc x̄',
         'Mean — trung bình',
         'M = 13,3 tuổi (V-NAMHS)'),
        ('Mdn',
         'Median — trung vị',
         'Dùng khi phân phối lệch'),
        ('Mode',
         'Yếu vị — giá trị xuất hiện nhiều nhất',
         'Ít dùng cho liên tục'),
        ('SD',
         'Standard Deviation — độ lệch chuẩn (đo spread)',
         'M = 13,3 ± 1,8 SD'),
        ('SE / SEM',
         'Standard Error of Mean — sai số chuẩn của mean. SE = SD / √n',
         'Nhỏ hơn SD. Dùng tính CI'),
        ('IQR',
         'Interquartile Range (Q3 – Q1) — 50 % dữ liệu ở giữa',
         'Dùng cho median trong non-parametric'),
        ('Min / Max',
         'Giá trị nhỏ nhất / lớn nhất',
         'Range = Max – Min'),
        ('%',
         'Percentage (tỷ lệ phần trăm)',
         '21,7 % VTN có vấn đề SKTT'),
    ])
P()

H2('E.2. EFFECT SIZE (Độ lớn hiệu ứng)')
MakeTable(
    ['Ký hiệu', 'Tên', 'Diễn giải', 'Ngưỡng'],
    [
        ('MD',
         'Mean Difference — khác biệt trung bình thô',
         'Dùng khi cùng thang. Đơn vị của thang.',
         'Không có universal'),
        ('SMD',
         'Standardized MD — chuẩn hoá',
         'MD / SD pooled. Không đơn vị. So được xuyên thang.',
         'Cohen: <0,2 trivial; 0,2 small; 0,5 medium; 0,8 large'),
        ('Cohen d',
         'Cohen\'s d — SMD gốc',
         'd = (M1 – M2) / SD_pooled. Dùng cho NC đơn lẻ.',
         'Giống SMD (đều Cohen cutoffs)'),
        ('Hedges g',
         'Hedges\' g — SMD hiệu chỉnh small-sample',
         'g = d × J(df). CHUẨN COCHRANE cho MA.',
         'Giống Cohen d'),
        ('OR',
         'Odds Ratio (tỷ số odds)',
         'odds(A) / odds(B). OR = 1: không effect; >1: tăng; <1: giảm.',
         'OR 1,5 small; 2,5 medium; 4 large (Chen 2010)'),
        ('RR',
         'Risk Ratio (tỷ số nguy cơ tương đối)',
         'P(bệnh | exposed) / P(bệnh | unexposed). Dễ diễn giải hơn OR khi tỷ lệ thấp.',
         'Tương tự OR'),
        ('HR',
         'Hazard Ratio (tỷ số nguy cơ tức thời)',
         'Dùng trong Cox regression / survival analysis.',
         'Tương tự OR'),
        ('β',
         'Beta — hệ số hồi quy',
         'Standardized: tăng 1 SD X → β SD Y. Unstandardized: tăng 1 unit X → β unit Y.',
         'β standardized: 0,1 small; 0,3 medium; 0,5 large'),
        ('r',
         'Pearson correlation',
         'Hệ số tương quan –1 đến +1.',
         'Cohen: 0,1 small; 0,3 medium; 0,5 large'),
        ('η²p (eta squared partial)',
         'Effect size ANOVA',
         'Tỷ lệ variance giải thích bởi factor.',
         '0,01 small; 0,06 medium; 0,14 large'),
        ('NNT',
         'Number Needed to Treat',
         'Số người cần can thiệp để 1 người có cải thiện.',
         'NNT 5–10 = tốt; 10–20 = acceptable cho prevention'),
        ('logOR',
         'Logarithm của OR',
         'Dùng trong meta-regression để normalise.',
         'exp(logOR) = OR. Ví dụ logOR 1,81 → OR ≈ 6,1'),
    ])
P()

H2('E.3. KHOẢNG TIN CẬY')
P('Xem chi tiết trong doc "Giải thích KTC vs CrI". Tóm tắt:', italic=True)
P('• 95 % CI = Confidence Interval tần suất (frequentist) — "nếu lặp lại NC vô hạn, 95 % CI sẽ chứa giá trị thật"', size=10)
P('• 95 % CrI = Credible Interval Bayesian — "có 95 % xác suất giá trị thật trong khoảng này"', size=10)
P('• 90 % (ít dùng) = hẹp hơn 95 %. Hay dùng trong equivalence testing (TOST)', size=10)
P()

H2('E.4. TEST STATISTICS')
MakeTable(
    ['Ký hiệu', 'Tên', 'Dùng khi', 'p-value significance'],
    [
        ('p',
         'p-value',
         'Mọi tests thống kê',
         '< 0,05 = có ý nghĩa thống kê; < 0,01 = highly significant'),
        ('t',
         't-statistic (Student t-test)',
         'So 2 means, n nhỏ, phân phối gần chuẩn',
         'Kèm df (degrees of freedom)'),
        ('F',
         'F-statistic (ANOVA)',
         'So ≥ 3 means hoặc test interaction',
         'Kèm df1, df2'),
        ('χ² (chi-square)',
         'Chi-square test',
         'So biến categorical (bảng 2×2, 2×k)',
         'Kèm df'),
        ('z',
         'z-statistic',
         'Mẫu lớn (CLT) hoặc tỷ lệ',
         'Z > 1,96 → p < 0,05'),
        ('r (Pearson) / ρ (Spearman)',
         'Correlation coefficient',
         'Đo độ tương quan',
         '|r| > 0,3 thường significant nếu n đủ'),
        ('Wilcoxon / Mann-Whitney U',
         'Non-parametric alternatives',
         'Khi không giả định phân phối chuẩn',
         'So median thay mean'),
    ])
P()

H2('E.5. META-ANALYSIS parameters')
MakeTable(
    ['Ký hiệu', 'Tên', 'Ý nghĩa', 'Ngưỡng'],
    [
        ('I²',
         'I-squared (heterogeneity)',
         'Tỷ lệ variance giữa các NC do heterogeneity (vs chance)',
         '0–25 % low; 25–75 % moderate; >75 % high'),
        ('τ² (tau-squared)',
         'Between-study variance',
         'Ước lượng variance thực giữa NC. Random-effects MA.',
         'Lớn = heterogeneity cao'),
        ('Q (Cochran\'s Q)',
         'Heterogeneity test',
         'Test giả thuyết heterogeneity = 0',
         'p < 0,05 → có heterogeneity'),
        ('SUCRA',
         'Surface Under Cumulative Ranking curve',
         'Xếp hạng can thiệp trong NMA. 0–100 %.',
         '> 70 % = xếp hạng cao'),
        ('P-score',
         'Frequentist alternative to SUCRA',
         'Tương tự SUCRA nhưng cho frequentist NMA',
         'Tương tự SUCRA'),
        ('Funnel plot',
         'Đồ thị publication bias',
         'Asymmetry = có khả năng publication bias',
         'Kèm Egger test'),
        ('Egger test',
         'Test asymmetry funnel plot',
         'p < 0,1 → có khả năng publication bias',
         '—'),
        ('Forest plot',
         'Đồ thị tóm tắt effect sizes',
         'Mỗi NC 1 dòng + pooled estimate dưới cùng',
         '—'),
    ])
P()

H2('E.6. BAYESIAN parameters')
MakeTable(
    ['Ký hiệu', 'Tên', 'Ý nghĩa'],
    [
        ('Posterior',
         'Phân phối hậu nghiệm',
         'Niềm tin cập nhật về tham số SAU khi thấy data. = Prior × Likelihood (luật Bayes).'),
        ('Prior',
         'Phân phối tiền nghiệm',
         'Niềm tin BAN ĐẦU (trước khi có data). Có thể "informative" (mạnh) hoặc "non-informative/vague" (yếu).'),
        ('Likelihood',
         'Hàm hợp lý (Likelihood function)',
         'Xác suất quan sát được data cho mỗi giá trị tham số.'),
        ('MCMC',
         'Markov Chain Monte Carlo',
         'Thuật toán sampling từ posterior khi không tính giải tích được. Output: 1.000–10.000 samples.'),
        ('Gibbs / Metropolis-Hastings / HMC / NUTS',
         'Các thuật toán MCMC cụ thể',
         'Stan dùng HMC/NUTS (hiệu quả nhất).'),
        ('R-hat (R̂)',
         'Potential Scale Reduction Factor',
         'Đo convergence MCMC. < 1,01 = converged tốt; > 1,1 = có vấn đề.'),
        ('ESS',
         'Effective Sample Size',
         'Số samples độc lập thực sự (sau khi loại autocorrelation). ESS > 1.000 = OK.'),
        ('Trace plot',
         'Đồ thị trace của MCMC',
         'Diagnostic visual — nếu chains mixing tốt, trace giống "fuzzy caterpillar".'),
        ('BF (Bayes Factor)',
         'Tỷ số bằng chứng cho 2 giả thuyết',
         'BF > 10 = strong evidence; BF > 30 = very strong; BF < 1/10 = evidence cho null.'),
        ('HDI / CrI',
         'Highest Density Interval / Credible Interval',
         'HDI = khoảng chứa highest density, có thể khác CrI central quantile.'),
    ])
P()
doc.add_page_break()

# ============================================================
# PHẦN F — KHÁI NIỆM LÂM SÀNG
# ============================================================
H1('PHẦN F. KHÁI NIỆM LÂM SÀNG')

H2('F.1. TỶ LỆ (Prevalence / Incidence)')
MakeTable(
    ['Loại', 'Định nghĩa', 'Ví dụ'],
    [
        ('Point prevalence',
         'Tỷ lệ HIỆN CÓ tại 1 thời điểm',
         'Lo âu hiện tại 10 %'),
        ('12-month prevalence',
         'Tỷ lệ đã CÓ trong 12 tháng qua',
         'V-NAMHS: 21,7 % VTN có vấn đề SKTT trong 12 tháng qua'),
        ('Lifetime prevalence',
         'Tỷ lệ đã TỪNG CÓ trong đời',
         'Kessler: 50 % người lớn Mỹ từng có MH disorder trong đời'),
        ('Incidence',
         'Tỷ lệ KHỞI PHÁT mới trong khoảng thời gian',
         'Ít phổ biến hơn prevalence trong tâm lý'),
        ('Weighted prevalence',
         'Tỷ lệ đã hiệu chỉnh theo survey weights (phân bố dân số)',
         'V-NAMHS dùng weighted theo tuổi–giới–đô thị/nông thôn'),
    ])
P()

H2('F.2. SUBTHRESHOLD vs FULL THRESHOLD')
note_box(
    'Đây là khái niệm CỐT LÕI trong NC tâm lý học hiện đại (V-NAMHS, QT049 Zhang...). Có 2 ngưỡng:\n\n'
    '• FULL THRESHOLD = đủ tiêu chí DSM-5 để chẩn đoán rối loạn. Kèm suy giảm chức năng.\n'
    '• SUBTHRESHOLD = đáp ứng ít nhất NỬA số triệu chứng — chưa đủ chẩn đoán nhưng đã có distress.\n\n'
    'V-NAMHS báo cáo CẢ HAI: "mental disorder" (full) = 3,3 %; "mental health problem" (subthreshold+) '
    '= 21,7 %. Ý nghĩa lâm sàng: subthreshold là nhóm QUAN TRỌNG cho can thiệp PHÒNG NGỪA trước khi '
    'tiến triển thành disorder đầy đủ.',
    color_hex='E3F2FD')
P()

H2('F.3. UNIVERSAL vs TARGETED vs INDICATED')
MakeTable(
    ['Loại phòng ngừa', 'Định nghĩa', 'Ví dụ'],
    [
        ('Universal',
         'Cho TẤT CẢ dân số — không chọn lọc',
         'CBT cho toàn bộ HS THCS (Happy House VN030)'),
        ('Selective',
         'Cho NHÓM NGUY CƠ CAO (ví dụ HS DTTS, LGBTQ)',
         'Chưa có NC VN nào ở đúng dạng này'),
        ('Targeted / Indicated',
         'Cho NGƯỜI có triệu chứng dưới ngưỡng (subthreshold)',
         'BESST UK (HS GAD-7 ≥ 5); QT042 Brown & Carter'),
        ('Treatment',
         'Cho người đã có rối loạn chẩn đoán',
         'Walkup 2008 CAMS; QT043 Bress Maya'),
    ])
note_box(
    'XU HƯỚNG: Zhang 2026 (Bayesian MA) cho thấy UNIVERSAL hiệu ứng NHỎ (SMD –0,19 lo âu). '
    'Brown & Carter 2025 (BESST) ủng hộ TARGETED. → Hiện tại literature đang chuyển hướng từ '
    'universal sang targeted để tối đa hoá cost-effectiveness.',
    color_hex='FFF3E0')
P()

H2('F.4. PROMOTION vs PREVENTION vs RESPONSE')
MakeTable(
    ['Cấp độ', 'Định nghĩa', 'Ví dụ'],
    [
        ('Promotion (nâng cao)',
         'Tăng cường sức khoẻ tâm thần cho ALL — kể cả người khoẻ mạnh',
         'Mindfulness trong giờ thể dục; Positive psychology curriculum'),
        ('Prevention (phòng ngừa)',
         'Giảm nguy cơ khởi phát rối loạn',
         'CBT universal/targeted; FRIENDS program'),
        ('Response / Treatment',
         'Điều trị khi đã có triệu chứng/rối loạn',
         'CAMS; Maya app; iCBT cho SAD'),
    ])
P('Menon 2025 (QT051) scoping review 69 NC: chỉ 3/69 = PROMOTION, 46/69 = PREVENTION, 23/69 = RESPONSE. '
  'Gap lớn ở PROMOTION.', italic=True, size=10)
P()

H2('F.5. IMPAIRMENT + DISTRESS')
P('• IMPAIRMENT = suy giảm chức năng (không làm được các hoạt động hàng ngày). '
  'DSM-5 YÊU CẦU impairment để chẩn đoán rối loạn. V-NAMHS đo 4 lĩnh vực: gia đình, bạn bè, trường/việc, '
  'căng thẳng cá nhân.', size=11)
P('• DISTRESS = đau khổ chủ quan. Có thể có mà không impairment.', size=11)
P()

H2('F.6. RISK vs PROTECTIVE FACTORS')
MakeTable(
    ['Loại yếu tố', 'Ví dụ từ thư viện (Chen 2025 COVID meta)'],
    [
        ('Risk factors (yếu tố nguy cơ)',
         'Nghiện Internet (logOR 1,81); Lo âu cha mẹ (1,06); Gánh nặng học đường (0,56); KT-XH gia đình thấp (–0,25)'),
        ('Protective factors (yếu tố bảo vệ)',
         'Chức năng cảm xúc tốt (–1,45); Hỗ trợ xã hội (–0,93); Gia đình kết nối; Chất lượng cuộc sống'),
    ])
P()
doc.add_page_break()

# ============================================================
# PHẦN G — CHUẨN BÁO CÁO
# ============================================================
H1('PHẦN G. CHUẨN BÁO CÁO KHOA HỌC')

H2('G.1. PRISMA 2020')
P('Preferred Reporting Items for Systematic Reviews and Meta-Analyses (Page et al. 2020, BMJ).', italic=True)
P('27-item checklist + flow diagram chuẩn. BẮT BUỘC cho mọi SR/MA chất lượng cao.', size=11)
P('Biến thể: PRISMA-ScR (Scoping Review), PRISMA-NMA (Network MA), PRISMA-IPD (Individual Patient Data).',
  size=11)
P('Papers trong thư viện dùng PRISMA: QT028, QT029, QT037, QT039, QT040, QT048, QT049, QT051.',
  size=10, italic=True)
P()

H2('G.2. GRADE (evidence certainty)')
P('Grading of Recommendations Assessment, Development, and Evaluation.', italic=True)
P('Đánh giá CHẤT LƯỢNG bằng chứng (evidence certainty) thành 4 mức:', bold=True)
P('• HIGH (⨁⨁⨁⨁) — rất tự tin', size=11)
P('• MODERATE (⨁⨁⨁◯) — khá tự tin', size=11)
P('• LOW (⨁⨁◯◯) — ít tự tin', size=11)
P('• VERY LOW (⨁◯◯◯) — không tự tin', size=11)
P('Hạ bậc vì: risk of bias, inconsistency, indirectness, imprecision, publication bias.', size=11)
P('Zhang 2026 (QT049) tự kết luận evidence ở mức low-moderate → khuyến nghị cần high-quality trials.',
  italic=True, size=10)
P()

H2('G.3. Cochrane RoB 2')
P('Cochrane Risk-of-Bias 2.0 (Sterne et al. 2019, BMJ).', italic=True)
P('Đánh giá risk of bias RCT theo 5 domain:', bold=True)
P('1. Randomization process', size=11)
P('2. Deviations from intended intervention', size=11)
P('3. Missing outcome data', size=11)
P('4. Measurement of outcome', size=11)
P('5. Selection of reported result', size=11)
P('Mỗi domain: Low / Some concerns / High. Overall: Low nếu tất cả low; High nếu ≥ 1 high.',
  size=11)
P()
P('Biến thể: ROBINS-I cho non-randomized; ROB-NRSI cho NRSI.', size=10)
P()

H2('G.4. STROBE, COREQ, CONSORT')
MakeTable(
    ['Chuẩn', 'Dùng cho', 'Papers dùng'],
    [
        ('STROBE',
         'Báo cáo NC quan sát (cross-sectional, cohort, case-control)',
         'VN001, VN002, VN022, VN029...'),
        ('CONSORT',
         'Báo cáo RCT',
         'QT038 De Silva, QT043 Bress, VN030 Happy House'),
        ('COREQ',
         'Báo cáo NC định tính (FGD, KII interview)',
         'VN022 UNICEF (FGD + KII)'),
        ('MOOSE',
         'Báo cáo MA của NC quan sát',
         'Ít dùng (PRISMA đã thay)'),
        ('TREND',
         'Báo cáo non-randomized interventions',
         'Có thể dùng cho VN030 (non-randomized cluster)'),
    ])
P()

H2('G.5. Đăng ký tiền khai (Pre-registration)')
MakeTable(
    ['Registry', 'Đối tượng', 'Papers dùng'],
    [
        ('PROSPERO',
         'Pre-registration cho SR/MA (York, UK)',
         'QT037 CRD42022336376; QT048 CRD42022316746; QT039'),
        ('ClinicalTrials.gov',
         'RCT — tiêu chuẩn cho tạp chí ICMJE',
         'QT043 Bress NCT05130246; VN030'),
        ('UMIN-CTR',
         'Japan trials registry',
         'QT045 Matsumoto UMIN000050064'),
        ('Open Science Framework (OSF)',
         'Pre-register rộng hơn (mọi NC)',
         'Ít dùng trong thư viện lõi'),
    ])
P()
doc.add_page_break()

# ============================================================
# PHẦN H — BẢNG ĐỐI CHIẾU NHANH
# ============================================================
H1('PHẦN H. BẢNG ĐỐI CHIẾU NHANH: Papers × Thuật ngữ chính')

H2('H.1. Các bài Việt Nam (VN001–VN030)')
MakeTable(
    ['Paper', 'Thiết kế', 'Công cụ đo chính', 'Kết quả nổi bật'],
    [
        ('VN001 Hoa 2024', 'Mixed-methods (CROSS + IDI)',
         'GAD-7', 'GAD-7 ≥ 5: 40,6 % HS THPT Hà Nội'),
        ('VN002 V-NAMHS 2022', 'Cross-sectional national survey',
         'DISC-5 (DSM-5 diagnostic)', '3,3 % rối loạn tâm thần; 21,7 % vấn đề SKTT'),
        ('VN020 Lộc 2024', 'Cross-sectional',
         'DASS-Y', 'Lo âu 50+ % HS THPT TPHCM'),
        ('VN022 UNICEF 2022', 'Mixed-methods',
         'SDQ-25, MDS3, ESSA, RCBI', '26,1 % rủi ro trung/cao SDQ; 52,2 % cyberbullying'),
        ('VN027 Dinh 2021', 'Cross-sectional',
         'SDQ-25', 'School factors analysis'),
        ('VN029 Duong 2025', 'Cross-sectional multi-center',
         'DASS-21 + YBRS', 'Lo âu 50,3 % TPHCM; SOMD–HRB link'),
        ('VN030 Happy House 2023', 'Cluster controlled (non-randomized)',
         'CESD-R, MHC-SF, CSES', 'd = 0,11 CBT universal; fade-out 6m'),
        ('Trần Nguyễn Ngọc 2018', 'Before-after (no control)',
         'HAM-A', 'Thư giãn – luyện tập cho RLLA người lớn'),
    ])
P()

H2('H.2. Các bài quốc tế (QT028–QT051 tiêu biểu)')
MakeTable(
    ['Paper', 'Thiết kế', 'Công cụ / Chuẩn', 'Kết quả'],
    [
        ('QT028 Zugman 2024', 'Review AJP',
         'DSM-5, SSRIs, CBT', 'Current + future pediatric anxiety tx'),
        ('QT029 Li 2025', 'Bayesian NMA 30 RCT',
         'PRISMA + SUCRA', 'CBT SUCRA 0,66; PE SUCRA 0,51'),
        ('QT037 Praptomojati 2024', 'Systematic review',
         'PRISMA, PROSPERO', '7 NC CA-CBT ĐNA; 0 từ VN'),
        ('QT038 De Silva 2024', 'Cluster RCT',
         'Cochrane RoB 2, CONSORT', 'β = –0,096, p = 0,038 Sri Lanka'),
        ('QT039 Xian 2024', 'Bayesian NMA 30 RCT',
         'PRISMA + SUCRA', 'iCBT SUCRA 71,2 % cho SAD'),
        ('QT040 Walder 2025', 'MA 21 RCT',
         'PRISMA', 'g = 0,878 guided DMHI SAD'),
        ('QT042 Brown & Carter 2025', 'Editorial',
         'Narrative', '4 mô hình UK: BESST, MHST, PLACES'),
        ('QT043 Bress 2024', 'RCT pilot',
         'CONSORT, ClinicalTrials', 'Cohen d = 1,04 Maya app'),
        ('QT045 Matsumoto 2024', 'RCT đa trung tâm',
         'LSAS-SR, UMIN-CTR', 'OR response 4,97 iCBT Japan'),
        ('QT048 Chen 2025', '3-level MA 141 NC',
         'PRISMA, PROSPERO', '29 risk + 14 protective factors COVID'),
        ('QT049 Zhang 2026', 'Bayesian MA 31 RCT',
         'PRISMA', 'SMD anxiety –0,19 (95 % CrI: –0,22; –0,17)'),
        ('QT051 Menon 2025', 'Scoping review',
         'PRISMA-ScR', '69 NC / 12 LMIC Đông Á-TBD'),
    ])
P()

# ============================================================
# PHỤ LỤC
# ============================================================
P('─' * 70, color=RGBColor(0x99, 0x99, 0x99))
H2('PHỤ LỤC: Tài liệu đi kèm')
P('Document này là PHẦN 3 trong bộ 3 tài liệu giải thích thuật ngữ:', italic=True)
P('• Phần 1: "Giai_thich_KTC_vs_CrI_cho_thay_v3.docx" — CI vs CrI', size=11)
P('• Phần 2: "Giai_thich_MA_vs_NMA_cho_thay.docx" — Meta-Analysis vs Network MA', size=11)
P('• Phần 3: "Giai_thich_toan_bo_thuat_ngu_CSDL.docx" (file này) — Từ điển đầy đủ', size=11, bold=True)
P()

H3('Nguồn tham khảo')
P('• APA Publication Manual (7th ed., 2020)', size=10)
P('• Cochrane Handbook for Systematic Reviews v6.4 (2023)', size=10)
P('• Borenstein et al. "Introduction to Meta-Analysis" (2nd ed., 2021)', size=10)
P('• DSM-5 (2013) + DSM-5-TR (2022)', size=10)
P('• ICD-11 (WHO 2022)', size=10)
P('• Page MJ et al. "PRISMA 2020" (BMJ 2021)', size=10)
P('• Sterne JAC et al. "RoB 2" (BMJ 2019)', size=10)
P('• GRADE Working Group (https://www.gradeworkinggroup.org/)', size=10)
P()

P('Tài liệu này do Nhóm dự án Lo âu biên soạn 15/04/2026. Mục đích: giúp người đọc báo cáo '
  'NC tâm lý học có thể tra cứu nhanh mọi thuật ngữ gặp trong 68 papers canonical của dự án.',
  italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80), align=WD_ALIGN_PARAGRAPH.CENTER)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT):,} bytes')
