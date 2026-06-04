# -*- coding: utf-8 -*-
"""Sinh hai phien ban mo rong Bai 2 (Can thiep) cho luan an tien si - 26/05/2026.
FULL  ~ 17-18k tu (7 muc lon: Mo dau / Khung ly thuyet / Phuong phap / Ket qua /
                  Ban luan / Ket luan / TLTK)
MEDIUM ~ 11-12k tu (5 muc lon: bo Khung ly thuyet, gon Phuong phap, gon Ban luan)
Chinh sua tu Bai2_CanThiep_HSTHCS_v6_16052026.docx + vac het loi audit 18/05.
Them 4 ref khung ly thuyet: Kendall 1994; Barrett-Dadds-Rapee 1996;
Mrazek-Haggerty 1994 (IOM tier); Bower-Gilbody 2005 (stepped care)."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUTDIR = os.path.join(ROOT, 'bai-bao-khgdvn')

# ============================================================
# HELPERS
# ============================================================
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.5)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h

def P(doc, text, indent=True):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    return p

def table_3col(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, htext in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = htext
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

def render(doc, blocks):
    for kind, *args in blocks:
        if kind == 'H1':
            H(doc, args[0], 1)
        elif kind == 'H2':
            H(doc, args[0], 2)
        elif kind == 'P':
            P(doc, args[0])
        elif kind == 'TABLE':
            headers, rows, widths = args
            table_3col(doc, headers, rows, widths)
        elif kind == 'CAP':
            p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(args[0]); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True

# ============================================================
# TLTK (21 entries: 17 cua v6 + 4 khung ly thuyet)
# ============================================================
TLTK = [
    "Barrett, P. M., Dadds, M. R., & Rapee, R. M. (1996). Family treatment of childhood anxiety: A controlled trial. Journal of Consulting and Clinical Psychology, 64(2), 333-342. https://doi.org/10.1037/0022-006X.64.2.333",
    "Beck, A. T. (1976). Cognitive therapy and the emotional disorders. International Universities Press.",
    "Bower, P., & Gilbody, S. (2005). Stepped care in psychological therapies: Access, effectiveness and efficiency. British Journal of Psychiatry, 186(1), 11-17. https://doi.org/10.1192/bjp.186.1.11",
    "Bradshaw, C. P., McDaniel, H. L., Pas, E. T., Debnam, K. J., Bottiani, J. H., Powell, N., Ialongo, N. S., Morgan-Lopez, A. A., & Lochman, J. E. (2025). Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology, 110, Article 101437. https://doi.org/10.1016/j.jsp.2025.101437",
    "Bress, J. N., Falk, A., Schier, M. M., Jaywant, A., Moroney, E., Dargis, M., Bennett, S. M., Scult, M. A., Volpp, K. G., Asch, D. A., Balachandran, M., Perlis, R. H., Lee, F. S., & Gunning, F. M. (2024). Efficacy of a mobile app-based intervention for young adults with anxiety disorders: A randomized clinical trial. JAMA Network Open, 7(8), Article e2428372. https://doi.org/10.1001/jamanetworkopen.2024.28372",
    "Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, Article 1594658. https://doi.org/10.3389/fpsyt.2025.1594658",
    "Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939-991.",
    "He, Q., Li, J., Wang, J., & Qu, Z. (2026). Preventing depression in Chinese children and adolescents: A pilot study of a brief school-based cognitive behavioral group program. Journal of Affective Disorders, 394, Article 120559. https://doi.org/10.1016/j.jad.2025.120559",
    "Kendall, P. C. (1994). Treating anxiety disorders in children: Results of a randomized clinical trial. Journal of Consulting and Clinical Psychology, 62(1), 100-110. https://doi.org/10.1037/0022-006X.62.1.100",
    "Li, S. H., Achilles, M. R., Spanos, S., Habak, S., Werner-Seidler, A., & ODea, B. (2022). A cognitive behavioural therapy smartphone app for adolescent depression and anxiety: Co-design of ClearlyMe. The Cognitive Behaviour Therapist, 15, Article e13. https://doi.org/10.1017/S1754470X22000095",
    "Liu, S., Xiao, H., Duan, Y., Shi, L., Wang, P., Cao, L., Li, H., Huang, X., & Qiu, C. (2025). CBT treatment delivery formats for generalized anxiety disorder: A systematic review and network meta-analysis of randomized controlled trials. Translational Psychiatry, 15, Article 197. https://doi.org/10.1038/s41398-025-03414-3",
    "Matsumoto, K., Hamatani, S., Shiga, K., Iiboshi, K., Kasai, M., Kimura, Y., Yokota, S., Watanabe, K., Kubo, Y., & Nakamura, M. (2024). Effectiveness of unguided internet-based cognitive behavioral therapy for subthreshold social anxiety disorder in adolescents and young adults: Multicenter randomized controlled trial. JMIR Pediatrics and Parenting, 7, Article e55786. https://doi.org/10.2196/55786",
    "Mrazek, P. J., & Haggerty, R. J. (Eds.). (1994). Reducing risks for mental disorders: Frontiers for preventive intervention research. National Academies Press. https://doi.org/10.17226/2139",
    "Praptomojati, A., Icanervilia, A. V., Nauta, M. H., & Bouman, T. K. (2024). A systematic review of culturally adapted cognitive behavioral therapy (CA-CBT) for anxiety disorders in Southeast Asia. Asian Journal of Psychiatry, 92, Article 103896. https://doi.org/10.1016/j.ajp.2023.103896",
    "Urao, Y., Yoshida, M., Koshiba, T., Sato, Y., Ishikawa, S., & Shimizu, E. (2018). Effectiveness of a cognitive behavioural therapy-based anxiety prevention programme at an elementary school in Japan: A quasi-experimental study. Child and Adolescent Psychiatry and Mental Health, 12, Article 33. https://doi.org/10.1186/s13034-018-0240-5",
    "Urao, Y., Yoshida, M., Sato, Y., & Shimizu, E. (2022). School-based cognitive behavioural intervention programme for addressing anxiety in 10- to 11-year-olds using short classroom activities in Japan: A quasi-experimental study. BMC Psychiatry, 22, Article 658. https://doi.org/10.1186/s12888-022-04326-y",
    "Walder, N., Frey, A., Berger, T., & Schmidt, S. J. (2025). Digital mental health interventions for the prevention and treatment of social anxiety disorder in children, adolescents, and young adults: A systematic review and meta-analysis of randomized controlled trials. Journal of Medical Internet Research, 27, Article e67067. https://doi.org/10.2196/67067",
    "Xian, J., Zhang, Y., & Jiang, B. (2024). Psychological interventions for social anxiety disorder in children and adolescents: A systematic review and network meta-analysis. Journal of Affective Disorders, 365, 614-627. https://doi.org/10.1016/j.jad.2024.08.097",
    "Tran Nguyen Ngoc. (2018). Danh gia hieu qua dieu tri roi loan lo au lan toa bang lieu phap thu gian luyen tap [Luan an tien si y hoc, Truong Dai hoc Y Ha Noi]. Truong Dai hoc Y Ha Noi.",
    "Tran, T. D., Nguyen, H., Shochet, I., Nguyen, N., La, N., Wurfl, A., Orr, J., Nguyen, H., Stocker, R., & Fisher, J. (2023). School-based universal mental health promotion intervention for adolescents in Vietnam: Two-arm, parallel, controlled trial. Cambridge Prisms: Global Mental Health, 10, Article e69. https://doi.org/10.1017/gmh.2023.66",
    "Vien Xa hoi hoc, Dai hoc Queensland, & Dai hoc Johns Hopkins. (2022). Viet Nam adolescent mental health survey (V-NAMHS): Report on main findings. Vien Xa hoi hoc.",
]

print(f"Tong so TLTK: {len(TLTK)}")

# ============================================================
# CONTENT BLOCKS
# ============================================================

SEC_MO_DAU = [
    ('H1', '1. Mở đầu chương'),
    ('P', 'Rối loạn lo âu ở học sinh trung học cơ sở (HSTHCS) là một vấn đề sức khỏe tâm thần đang được ghi nhận với tỷ lệ ngày càng cao tại Việt Nam và trên thế giới. Trên cơ sở kết quả tổng hợp về yếu tố nguy cơ đã trình bày ở chương trước, chương này tập trung vào câu hỏi: với mỗi loại nguy cơ đã được nhận diện, hiện có những can thiệp tâm lý nào đã được kiểm định bằng bằng chứng thực nghiệm, và những hình thức can thiệp nào khả thi trong bối cảnh giáo dục và hệ thống chăm sóc sức khỏe tâm thần của Việt Nam? Trả lời câu hỏi này là bước chuyển từ "biết nguy cơ" sang "biết cách giảm nguy cơ", và là nền tảng cho phần thực trạng và thực nghiệm ở các chương sau của luận án.'),
    ('P', 'Liệu pháp nhận thức – hành vi (Cognitive Behavioral Therapy – CBT) thường được coi là tiếp cận chuẩn vàng cho điều trị rối loạn lo âu ở trẻ em – vị thành niên. Phát triển từ nền tảng lý thuyết của Beck (1976) và được kiểm định hiệu quả trên trẻ em từ thử nghiệm lâm sàng ngẫu nhiên đầu tiên của Kendall (1994), CBT đã trải qua hơn bốn thập kỷ thử nghiệm thực nghiệm và hiện được phân nhánh thành nhiều phiên bản giao thức khác nhau. Các giao thức tiêu biểu gồm Coping Cat của Kendall ở Hoa Kỳ, FRIENDS phát triển từ công trình của Barrett, Dadds và Rapee (1996) tại Úc và sau đó tiếp nối bởi Cool Kids, và các bản địa hóa cho khu vực châu Á như Journey of the Brave tại Nhật Bản (Urao và cộng sự, 2018; Urao và cộng sự, 2022).'),
    ('P', 'Phân tích tổng hợp lớn của Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành niên xác nhận chiến lược đối phó tập trung vào vấn đề – một thành tố cốt lõi của CBT – có liên hệ ngược chiều và đáng kể với triệu chứng lo âu. Tổng quan hệ thống và phân tích mạng lưới của Liu và cộng sự (2025) trên 52 thử nghiệm lâm sàng ngẫu nhiên ở bệnh nhân rối loạn lo âu tổng quát — chủ yếu là người trưởng thành — cho thấy CBT cá nhân được xếp hạng hiệu quả cao nhất trong số các hình thức cung cấp. Phân tích mạng lưới riêng cho rối loạn lo âu xã hội ở trẻ em – vị thành niên của Xian và cộng sự (2024) báo cáo kết quả cùng hướng, trong đó các hình thức CBT — đặc biệt là CBT qua internet — được xếp hạng hiệu quả cao nhất. Trên phương diện học đường, phân tích tổng hợp của Cai và cộng sự (2025) ghi nhận các chương trình tại trường có hiệu ứng tổng hợp nhỏ trên năng lực phục hồi tâm lý — một cấu trúc bảo vệ có liên quan đến việc giảm nguy cơ lo âu.'),
    ('P', 'Trong bối cảnh chuyển đổi số mạnh mẽ và sự gia tăng tỷ lệ sở hữu điện thoại thông minh ở vị thành niên, can thiệp CBT qua nền tảng số (Digital Mental Health Intervention – DMHI) ngày càng được chú ý. Phân tích tổng hợp của Walder và cộng sự (2025) trên 21 thử nghiệm lâm sàng ngẫu nhiên về can thiệp số cho rối loạn lo âu xã hội ở trẻ em, vị thành niên và thanh niên trẻ báo cáo hiệu ứng tổng hợp g = 0,508, trong đó nhóm can thiệp số có hướng dẫn đạt g = 0,825. Tại Việt Nam, các nghiên cứu can thiệp dựa trên CBT hiện có chủ yếu mang tính thử nghiệm pilot quy mô nhỏ hoặc nhắm vào nhóm tuổi khác; các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên HSTHCS Việt Nam vẫn chưa được triển khai.'),
    ('P', 'Câu hỏi đặt ra cho nghiên cứu can thiệp Việt Nam hiện nay không phải là CBT có hiệu quả hay không – câu hỏi này đã được trả lời rõ qua các phân tích tổng hợp quốc tế. Câu hỏi quan trọng hơn là làm thế nào để bản địa hóa CBT cho bối cảnh giáo dục Việt Nam, triển khai khả thi trong điều kiện thiếu hụt nhân lực tâm lý học đường, và tích hợp can thiệp số với hệ thống chăm sóc truyền thống. Chương này hệ thống hóa các nghiên cứu công bố giai đoạn 2015–2026 theo ba nhóm: (i) CBT đã xác nhận hiệu quả qua thử nghiệm lâm sàng, (ii) CBT triển khai tại trường, và (iii) CBT qua nền tảng số. Bên cạnh đó, chương cũng thảo luận về các mô hình kết hợp như chăm sóc theo bậc (stepped care) và can thiệp pha trộn (blended) — những mô hình có tiềm năng giải quyết bài toán phân tầng dịch vụ trong điều kiện nguồn lực hạn chế của Việt Nam.'),
    ('P', 'Mục tiêu của chương là đặt nền tảng tổng quan cho các thiết kế nghiên cứu can thiệp ưu tiên tại Việt Nam giai đoạn 2026–2030. Tổng quan này, kết hợp với chương về yếu tố nguy cơ, tạo thành một cặp tham chiếu hai chiều: yếu tố nguy cơ chỉ ra "đâu là điểm cần tác động", còn các can thiệp hiệu quả chỉ ra "tác động bằng cách nào". Sự ghép nối này là điều kiện cần để các chương thực trạng và thực nghiệm tiếp theo có thể đề xuất một mô hình can thiệp phù hợp với điều kiện văn hóa, giáo dục và y tế của Việt Nam.'),
]

SEC_KHUNG = [
    ('H1', '2. Khung lý thuyết phân tích can thiệp tâm lý'),
    ('P', 'Phân tích các nghiên cứu can thiệp đòi hỏi một khung lý thuyết kép: một khung về cơ chế can thiệp (CBT làm gì và tác động qua kênh nào) và một khung về cấp độ dịch vụ (can thiệp được tổ chức theo tầng nào trong hệ thống chăm sóc). Mục này trình bày ba khung lý thuyết được sử dụng xuyên suốt chương: khung sức khỏe tâm thần cộng đồng ba tầng của Mrazek và Haggerty (1994), khung CBT cho trẻ em – vị thành niên xuất phát từ Beck (1976) và được mở rộng bởi Kendall (1994) cùng Barrett, Dadds, Rapee (1996), và khung tích hợp ba kênh can thiệp đề xuất cho HSTHCS Việt Nam.'),

    ('H2', '2.1. Khung sức khỏe tâm thần cộng đồng ba tầng'),
    ('P', 'Báo cáo của Viện Y học Hoa Kỳ do Mrazek và Haggerty (1994) chủ biên đề xuất khung phân tầng can thiệp phòng ngừa thành ba loại: phổ quát (universal), chọn lọc (selective), và chỉ định (indicated). Can thiệp phổ quát hướng tới toàn bộ dân số mục tiêu mà không qua sàng lọc — ví dụ tích hợp giáo dục cảm xúc – xã hội vào chương trình môn học chính khóa. Can thiệp chọn lọc hướng tới nhóm có nguy cơ cao nhưng chưa có triệu chứng — ví dụ học sinh sống trong gia đình có cha mẹ ly hôn hoặc có anh chị em mắc rối loạn lo âu. Can thiệp chỉ định hướng tới các cá nhân đã có triệu chứng cận lâm sàng — ví dụ học sinh được sàng lọc dương tính với thang đo lo âu nhưng chưa đạt ngưỡng chẩn đoán lâm sàng. Mỗi tầng có triết lý, mục tiêu kết cục và yêu cầu nguồn lực khác nhau.'),
    ('P', 'Khung ba tầng có hai hệ quả quan trọng cho phân tích chương này. Thứ nhất, các phân tích tổng hợp như của Cai và cộng sự (2025) về can thiệp tại trường thường gộp cả ba tầng vào một ước tính hiệu ứng tổng hợp, dẫn tới hiệu ứng trung bình bị kéo nhỏ — tầng phổ quát có cỡ hiệu ứng thấp hơn vì áp dụng cho cả những học sinh không có vấn đề. Đánh giá đúng hiệu quả cần phân tách theo tầng. Thứ hai, một hệ thống can thiệp khả thi cho Việt Nam phải có cả ba tầng: tầng phổ quát chi phí thấp tích hợp vào giáo dục công dân hoặc hoạt động trải nghiệm; tầng chọn lọc dành cho nhóm rủi ro cao như học sinh chuyển trường, học sinh có rối loạn học tập; và tầng chỉ định dành cho học sinh có triệu chứng cận lâm sàng được phát hiện qua sàng lọc trường học.'),

    ('H2', '2.2. Khung CBT cho trẻ em – vị thành niên'),
    ('P', 'Mô hình CBT cổ điển của Beck (1976) đặt nền tảng trên giả định rằng các mô hình nhận thức không thích nghi (cognitive distortions) duy trì các phản ứng cảm xúc tiêu cực; can thiệp tập trung vào nhận diện, thử nghiệm và điều chỉnh các suy nghĩ tự động cùng các niềm tin cốt lõi. Đối với trẻ em – vị thành niên, do năng lực nhận thức về siêu nhận thức (metacognition) chưa hoàn thiện, các giao thức gốc của Beck cần được mở rộng để tích hợp thêm các thành phần hành vi như tiếp xúc dần (graded exposure), kỹ thuật thư giãn, và các bài tập tăng cường nhận thức cảm xúc.'),
    ('P', 'Thử nghiệm lâm sàng ngẫu nhiên đầu tiên về CBT cho trẻ em mắc rối loạn lo âu là của Kendall (1994) – công trình giới thiệu chương trình Coping Cat, một giao thức 16 buổi cho trẻ 9–13 tuổi mắc rối loạn lo âu chia ly, rối loạn lo âu tổng quát hoặc rối loạn lo âu xã hội. Coping Cat sử dụng cấu trúc FEAR — viết tắt của Feeling frightened, Expecting bad things, Attitudes and actions, Results and rewards — như một công cụ ghi nhớ giúp trẻ tự áp dụng các bước CBT. Ngay sau đó, Barrett, Dadds và Rapee (1996) tại Úc phát triển chương trình can thiệp gia đình (FAM-CBT) — về sau trở thành nền tảng cho các giao thức như FRIENDS và Cool Kids, bổ sung thành phần gia đình vào giao thức và mở rộng đối tượng sang nhóm vị thành niên. Hai chương trình này, cùng với các bản kế thừa của chúng ở các quốc gia khác, đã được kiểm chứng qua hàng trăm thử nghiệm lâm sàng ngẫu nhiên trong ba thập kỷ qua.'),
    ('P', 'Bản địa hóa CBT cho từng nền văn hóa là vấn đề liên tục được thảo luận. Praptomojati và cộng sự (2024) trong tổng quan hệ thống các phiên bản CBT thích nghi văn hóa (CA-CBT) cho rối loạn lo âu tại Đông Nam Á nhận định rằng các yếu tố như tôn giáo, gia đình mở rộng và quan niệm cộng đồng về sức khỏe tâm thần cần được tích hợp tường minh vào giao thức nhằm nâng cao hiệu quả ở khu vực này. Đối với Việt Nam, ba yếu tố văn hóa đặc biệt cần lưu ý là: vai trò của gia đình mở rộng (ông bà, cô dì) trong nuôi dạy con cái; kỳ thị xã hội đối với "bệnh tâm thần" vẫn còn cao; và niềm tin tâm linh có thể đóng vai trò vừa bảo vệ vừa cản trở khi định hướng tìm kiếm hỗ trợ chuyên môn.'),

    ('H2', '2.3. Khung tích hợp ba kênh can thiệp cho HSTHCS Việt Nam'),
    ('P', 'Trên cơ sở khung ba tầng của Mrazek-Haggerty và khung CBT đã trình bày, luận án đề xuất một khung phân tích tích hợp với ba kênh cung cấp can thiệp: (i) kênh lâm sàng — CBT cá nhân hoặc nhóm thực hiện bởi nhà tâm lý lâm sàng tại bệnh viện hoặc phòng tham vấn chuyên khoa; (ii) kênh trường học — CBT triển khai bởi giáo viên tâm lý học đường hoặc giáo viên chủ nhiệm được đào tạo, lồng ghép vào chương trình môn học hoặc hoạt động trải nghiệm; (iii) kênh số — CBT cung cấp qua ứng dụng điện thoại, nền tảng internet hoặc mô hình pha trộn (blended). Ba kênh này không loại trừ lẫn nhau mà có thể vận hành đồng thời và bổ sung cho nhau trong một hệ thống chăm sóc đa tầng.'),
    ('P', 'Mỗi kênh có ưu thế và giới hạn riêng. Kênh lâm sàng có chất lượng chuyên môn cao nhất, phù hợp cho học sinh có triệu chứng nặng hoặc có rối loạn đồng diễn, nhưng có khả năng tiếp cận thấp do thiếu nhân lực chuyên khoa và chi phí cao. Kênh trường học có tiềm năng tiếp cận lớn nhất do trường học là nơi học sinh đến hằng ngày, đồng thời giảm rào cản kỳ thị và chi phí; tuy nhiên chất lượng can thiệp phụ thuộc vào năng lực và thời gian của giáo viên. Kênh số có ưu thế về quy mô và linh hoạt thời gian, nhưng phụ thuộc vào điều kiện thiết bị, kết nối internet và động lực tự thân của học sinh để duy trì sử dụng.'),
    ('P', 'Khung tích hợp gợi ý một mô hình chăm sóc theo bậc (stepped care) cho Việt Nam, dựa trên nguyên lý của Bower và Gilbody (2005): học sinh được phân bậc tiếp cận theo mức độ triệu chứng và đáp ứng. Bậc 1 (mức độ cận lâm sàng nhẹ): can thiệp số tự hướng dẫn hoặc tài liệu tự giúp; nếu không cải thiện sau 4–6 tuần, chuyển lên bậc 2. Bậc 2 (mức độ cận lâm sàng vừa): CBT nhóm tại trường do giáo viên tâm lý dẫn dắt, hoặc can thiệp số có hướng dẫn; nếu không cải thiện hoặc triệu chứng nặng lên, chuyển lên bậc 3. Bậc 3 (mức độ lâm sàng): CBT cá nhân với nhà tâm lý lâm sàng, kết hợp tâm thần khoa khi cần thiết. Mô hình stepped care cho phép sử dụng hiệu quả nguồn lực chuyên môn hạn chế bằng cách dành can thiệp chuyên sâu cho những trường hợp thực sự cần.'),
    ('CAP', 'Bảng 1. Khung tích hợp ba kênh × ba tầng can thiệp cho HSTHCS Việt Nam'),
    ('TABLE', ['Tầng / Kênh', 'Kênh lâm sàng', 'Kênh trường học', 'Kênh số'],
              [['Phổ quát', 'Không phù hợp', 'Giáo dục cảm xúc – xã hội lồng ghép môn học', 'Ứng dụng tự học miễn phí'],
               ['Chọn lọc', 'Tham vấn ngắn 4–6 buổi cho nhóm rủi ro', 'CBT nhóm 6–8 buổi cho học sinh sàng lọc dương tính', 'Can thiệp số có hướng dẫn từ xa'],
               ['Chỉ định', 'CBT cá nhân 12–16 buổi với chuyên gia', 'Tham vấn cá nhân tại trường + chuyển tuyến', 'Blended: ứng dụng + tham vấn từ xa định kỳ']],
              [3.5, 4.0, 4.0, 4.0]),

    ('H2', '2.4. Phê phán và giới hạn của khung tích hợp'),
    ('P', 'Khung tích hợp đề xuất có ba giới hạn cần lưu ý. Thứ nhất, khung dựa trên giả định rằng cả ba kênh có thể triển khai song song, trong khi thực tế Việt Nam hiện thiếu trầm trọng nhân lực kênh lâm sàng và chưa có nền tảng kênh số tiếng Việt được kiểm định cho HSTHCS. Việc xây dựng đồng bộ ba kênh sẽ đòi hỏi thời gian và đầu tư đáng kể. Thứ hai, mô hình stepped care giả định tồn tại một hệ thống sàng lọc – chuyển tuyến vận hành ổn định, nhưng hiện chưa có quy trình chuẩn nào về sàng lọc lo âu cho học sinh THCS trong hệ thống giáo dục Việt Nam. Thứ ba, khung này chủ yếu dựa trên các nghiên cứu can thiệp ở các nước thu nhập cao và Trung Quốc, có thể không phản ánh hết đặc thù của khu vực nông thôn, miền núi và dân tộc thiểu số Việt Nam.'),
    ('P', 'Mặc dù có những giới hạn nêu trên, khung tích hợp vẫn cung cấp một công cụ phân tích hữu ích để: (i) phân loại các nghiên cứu hiện có theo kênh cung cấp, làm rõ chỗ nào đã có bằng chứng và chỗ nào còn khoảng trống; (ii) đề xuất các thiết kế nghiên cứu can thiệp tương lai theo hướng tích hợp ba kênh thay vì chỉ chọn một; (iii) định hướng các khuyến nghị chính sách về phân bổ nguồn lực giữa hệ thống chuyên khoa, hệ thống giáo dục và phát triển công nghệ.'),
]

SEC_PHUONG_PHAP = [
    ('H1', '3. Phương pháp tổng quan tài liệu'),

    ('H2', '3.1. Cách tiếp cận'),
    ('P', 'Chương này áp dụng thiết kế tổng quan tự sự (narrative review) thay vì tổng quan hệ thống chính thức (systematic review). Lựa chọn thiết kế dựa trên ba lý do: (i) phạm vi của tổng quan trải dài ba kênh can thiệp với các thiết kế nghiên cứu khác nhau (RCT, tựa thực nghiệm, pilot, phân tích tổng hợp), cần một khung tổng hợp linh hoạt đủ để gộp các dạng bằng chứng này; (ii) mục tiêu của tổng quan là nhận diện khoảng trống và đề xuất hướng nghiên cứu tương lai cho Việt Nam, hơn là ước tính một hiệu ứng tổng hợp duy nhất; (iii) số lượng nghiên cứu can thiệp về CBT cho rối loạn lo âu trong bối cảnh Việt Nam còn quá ít để áp dụng phân tích tổng hợp định lượng. Mặc dù vậy, để duy trì sự nghiêm túc của khâu chọn lọc và đánh giá, tổng quan tự sự ở đây áp dụng một số nguyên tắc của tổng quan hệ thống về tiêu chí lựa chọn, đối chiếu nguồn gốc và đánh giá chất lượng nghiên cứu.'),

    ('H2', '3.2. Tiêu chí lựa chọn tài liệu'),
    ('P', 'Khung lựa chọn tài liệu được xác lập trên ba tiêu chí. Thứ nhất về phạm vi thời gian, các công trình được ưu tiên là những bài công bố từ 2015 đến 2026; nhóm tác giả đặc biệt nhấn mạnh giai đoạn sau 2020 – thời điểm chứng kiến bước phát triển nhanh của can thiệp CBT số trong và sau đại dịch COVID-19. Một số tài liệu kinh điển trước 2015 vẫn được giữ lại khi cung cấp khung lý thuyết hoặc giao thức gốc, ví dụ công trình của Beck (1976) cho khung CBT, Kendall (1994) cho thử nghiệm lâm sàng ngẫu nhiên đầu tiên về CBT trên trẻ em mắc rối loạn lo âu, Barrett-Dadds-Rapee (1996) cho can thiệp gia đình, hay khung phòng ngừa ba tầng của Mrazek và Haggerty (1994).'),
    ('P', 'Thứ hai về đối tượng, trọng tâm là HSTHCS Việt Nam (11–15 tuổi). Khi gặp các thử nghiệm có mẫu hỗn hợp tuổi (trẻ em – vị thành niên), nhóm tác giả trích xuất riêng dữ liệu khối tuổi đầu vị thành niên hoặc thông số chung khi không tách rời được khối tuổi. Các nghiên cứu trên mẫu trẻ tiền học đường hoặc thanh niên trên 18 tuổi không được đưa vào trừ khi là phân tích tổng hợp đa nhóm tuổi có giá trị khung lý thuyết — ví dụ Bress và cộng sự (2024) có mẫu 18–25 tuổi vẫn được giữ lại vì cung cấp bằng chứng về hiệu quả của một ứng dụng CBT số được thiết kế chỉn chu, thông tin có ý nghĩa khi suy luận cho thiết kế can thiệp số cho lứa tuổi nhỏ hơn.'),
    ('P', 'Thứ ba về nội dung, các nghiên cứu được chọn là thử nghiệm lâm sàng ngẫu nhiên, thử nghiệm bán ngẫu nhiên, nghiên cứu pilot có nhóm chứng, hoặc phân tích tổng hợp trên các thử nghiệm này. Kết cục được đo bằng các thang lo âu chuẩn hóa cho trẻ em – vị thành niên (RCADS, SCAS, MASC, SCARED) hoặc các thước đo tổng (DASS-21, STAI). Các bài bình luận, ý kiến chuyên gia chưa qua bình duyệt cùng luận án tiến sĩ chưa bảo vệ không nằm trong danh sách trích dẫn của bài viết, ngoại trừ luận án tiến sĩ y học của Trần Nguyễn Ngọc (2018) — đã bảo vệ và là nguồn dữ liệu trong nước hiếm hoi về liệu pháp thư giãn trong điều trị rối loạn lo âu lan tỏa.'),

    ('H2', '3.3. Nguồn dữ liệu'),
    ('P', 'Nguồn tài liệu phân tích trong chương được rút từ thư mục nghiên cứu tích lũy trong quá trình thực hiện luận án của tác giả — gồm 24 công trình tiếng Việt và 73 công trình quốc tế (tổng 97 tài liệu), được phân loại theo chủ đề can thiệp (CBT lâm sàng, CBT trường học, CBT số) và theo khu vực địa lý. Các công trình quốc tế được trích từ các nền tảng cơ sở dữ liệu điện tử PubMed, Web of Science, Scopus, PsycINFO và các bản truy cập mở của Lancet, BMC, Springer, Frontiers, JMIR, Wiley, JAMA Network. Các công trình tiếng Việt được tra cứu trong Tạp chí Y học Việt Nam, Tạp chí Tâm lý học, Tạp chí Khoa học Giáo dục Việt Nam cùng các kỷ yếu hội thảo quốc gia về sức khỏe tâm thần học đường.'),
    ('P', 'Các từ khóa tìm kiếm tiếng Anh sử dụng là tổ hợp giữa nhóm thuật ngữ về can thiệp (cognitive behavioral therapy, CBT, school-based intervention, digital mental health intervention, smartphone app) và nhóm thuật ngữ về đối tượng (adolescent, child, young people, secondary school student, middle school student) cùng nhóm thuật ngữ về rối loạn (anxiety disorder, social anxiety, generalized anxiety, separation anxiety). Các từ khóa tiếng Việt tương ứng được sử dụng để tra cứu các nguồn trong nước.'),

    ('H2', '3.4. Quy trình trích xuất và thẩm định'),
    ('P', 'Quy trình trích xuất tài liệu gồm bốn bước. Bước 1: rà soát tiêu đề và tóm tắt để loại bỏ các bài không liên quan. Bước 2: đọc toàn văn các bài đạt sàng lọc sơ bộ và trích xuất các thông tin định lượng quan trọng (cỡ mẫu, đặc điểm đối tượng, thiết kế nghiên cứu, can thiệp, thước đo kết cục, hiệu ứng tổng hợp). Bước 3: đối chiếu mỗi thông số định lượng nêu trong bản thảo với báo cáo gốc để bảo đảm tính chính xác. Bước 4: đánh giá chất lượng nghiên cứu theo các tiêu chí phù hợp với thiết kế (RoB 2 cho RCT, ROBINS-I cho nghiên cứu tựa thực nghiệm, AMSTAR-2 cho phân tích tổng hợp).'),
    ('P', 'Trong khâu thực hiện bài viết, nhóm tác giả có dùng một mô hình ngôn ngữ lớn ở vai trò hỗ trợ kỹ thuật — tìm kiếm sơ bộ, kiểm tra chuẩn trích dẫn và rà soát hành văn. Toàn bộ phần đánh giá bằng chứng giữa các giao thức can thiệp, diễn giải kết quả và đề xuất hướng triển khai thuộc về trách nhiệm chuyên môn của nhóm. Mỗi thông số định lượng nêu trong bản thảo đều được đối chiếu lại với báo cáo gốc — quy trình tuân theo chính sách liêm chính nghiên cứu khi ứng dụng AI mà các tạp chí khoa học giáo dục Việt Nam đã công bố, đồng thời tham chiếu các nguyên tắc của Đạo luật AI châu Âu.'),

    ('H2', '3.5. Tổng hợp định tính và định lượng'),
    ('P', 'Do thiết kế tổng quan tự sự, tổng hợp định lượng (meta-analysis) không được thực hiện trong chương này. Thay vào đó, các hiệu ứng tổng hợp được trích xuất trực tiếp từ các phân tích tổng hợp đã công bố — ví dụ g = 0,508 của Walder và cộng sự (2025), SMD = 0,17 của Cai và cộng sự (2025), Cohen d = 0,93 đến 1,07 của Bress và cộng sự (2024). Các con số này được sử dụng để so sánh tương đối giữa các kênh can thiệp và giữa các phân nhóm trong cùng một kênh (ví dụ can thiệp số có hướng dẫn so với không hướng dẫn).'),
    ('P', 'Tổng hợp định tính được thực hiện theo phương pháp narrative synthesis: nhóm tác giả sắp xếp các nghiên cứu theo ba kênh, trong mỗi kênh phân theo nhóm dân số mục tiêu (vị thành niên đầu, vị thành niên giữa/cuối) và theo dạng rối loạn (lo âu tổng quát, lo âu xã hội, lo âu chia ly). Các phát hiện chung được tổng hợp theo bốn trục: hiệu quả, tính khả thi, độ tuân thủ và độ bền vững.'),

    ('H2', '3.6. Hạn chế của thiết kế tổng quan tự sự'),
    ('P', 'Thiết kế tổng quan tự sự có ba hạn chế cần ghi nhận. Thứ nhất, không kiểm soát được hoàn toàn rủi ro chọn lọc thiên kiến trong khâu chọn tài liệu — mặc dù đã áp dụng tiêu chí rõ ràng, việc không thực hiện hai khâu sàng lọc độc lập có thể dẫn tới sai lệch nhỏ. Thứ hai, không cung cấp được ước tính hiệu ứng tổng hợp định lượng cho riêng bối cảnh Việt Nam do số lượng nghiên cứu trong nước quá ít. Thứ ba, các phân tích so sánh chéo giữa các kênh can thiệp mang tính định tính và cần được xác nhận lại qua các thử nghiệm so sánh trực tiếp trong tương lai. Các hạn chế này được xem xét khi rút ra các kết luận và khuyến nghị ở mục Bàn luận.'),
]

SEC_KET_QUA = [
    ('H1', '4. Kết quả nghiên cứu'),

    ('H2', '4.1. CBT đã được xác nhận hiệu quả qua thử nghiệm lâm sàng'),
    ('P', 'CBT đã được kiểm định qua hơn bốn thập kỷ thử nghiệm lâm sàng trên trẻ em – vị thành niên với rối loạn lo âu. Khung CBT cổ điển được Beck (1976) đề xuất dựa trên giả định rằng các mô hình nhận thức không thích nghi duy trì các phản ứng cảm xúc tiêu cực; do đó can thiệp tập trung vào nhận diện, thử nghiệm và điều chỉnh các suy nghĩ tự động cùng các niềm tin cốt lõi. Đối với trẻ em – vị thành niên, các giao thức CBT thường tích hợp thêm các thành phần hành vi như tiếp xúc dần (graded exposure), kỹ thuật thư giãn và tăng cường nhận thức cảm xúc. Thử nghiệm lâm sàng ngẫu nhiên đầu tiên về CBT cho trẻ em mắc rối loạn lo âu của Kendall (1994) báo cáo tỷ lệ cải thiện 64% ở nhóm điều trị so với 5% ở nhóm chờ điều trị, đặt nền tảng cho hàng trăm thử nghiệm tiếp theo.'),
    ('P', 'Bằng chứng về hiệu quả so sánh giữa các hình thức cung cấp CBT chủ yếu đến từ nghiên cứu trên người trưởng thành. Tổng quan hệ thống và phân tích mạng lưới của Liu và cộng sự (2025) trên 52 thử nghiệm lâm sàng ngẫu nhiên ở bệnh nhân rối loạn lo âu tổng quát cho thấy CBT cá nhân đạt hiệu quả cao nhất, vượt trội so với CBT từ xa, điều trị thông thường và nhóm chứng đợi; CBT nhóm cũng vượt trội so với nhóm chứng đợi, trong khi CBT từ xa không cho thấy ưu thế rõ so với điều trị thông thường hay nhóm chứng đợi. Mặc dù mẫu nghiên cứu là người trưởng thành (tuổi trung bình khoảng 43), thứ bậc hiệu quả này là một tham chiếu hữu ích khi lựa chọn hình thức triển khai CBT cho học sinh, đồng thời cho thấy hiệu quả của một giao thức phụ thuộc đáng kể vào việc nó được thực hiện đầy đủ và có hỗ trợ chuyên môn hay không.'),
    ('P', 'Cần lưu ý rằng các phát hiện từ phân tích mạng lưới trên người trưởng thành không nên được áp dụng cơ học cho lứa tuổi vị thành niên. Ở vị thành niên, một số yếu tố có thể đảo ngược thứ bậc — ví dụ CBT nhóm có thể có lợi thế bổ sung so với CBT cá nhân do tận dụng được hiệu ứng nhóm đồng đẳng, một yếu tố phát triển quan trọng ở lứa tuổi này. Hơn nữa, CBT cho vị thành niên thường có thêm thành phần làm việc với cha mẹ và giáo viên, các thành phần này không có hoặc rất hạn chế trong CBT cho người trưởng thành. Do vậy, các kết quả của Liu và cộng sự (2025) cần được diễn giải thận trọng khi áp dụng cho thiết kế can thiệp HSTHCS Việt Nam, và lý tưởng nhất là được bổ sung bằng các phân tích mạng lưới riêng cho lứa tuổi vị thành niên trong tương lai.'),
    ('P', 'Đối với rối loạn lo âu xã hội – một trong những dạng phổ biến ở giai đoạn đầu vị thành niên – phân tích mạng lưới của Xian và cộng sự (2024) trên 30 thử nghiệm lâm sàng ngẫu nhiên ở trẻ em và vị thành niên xác định các hình thức CBT chiếm những vị trí xếp hạng hiệu quả cao nhất, trong đó CBT qua internet đứng đầu (chỉ số diện tích dưới đường xếp hạng tích lũy (SUCRA) 71,2%), tiếp đến là CBT nhóm và CBT cá nhân. Praptomojati và cộng sự (2024) trong tổng quan hệ thống các phiên bản CBT thích nghi văn hóa (CA-CBT) cho rối loạn lo âu tại Đông Nam Á nhận định rằng các yếu tố như tôn giáo, gia đình mở rộng và quan niệm cộng đồng về sức khỏe tâm thần cần được tích hợp tường minh vào giao thức nhằm nâng cao hiệu quả ở khu vực này. Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành niên xác nhận rằng đối phó tập trung vào vấn đề và đánh giá lại nhận thức có liên hệ ngược chiều và bền vững với triệu chứng lo âu – trầm cảm, củng cố cho nền tảng lý thuyết của CBT.'),
    ('P', 'Hai giao thức CBT nền tảng cần được nhắc đến chi tiết. Chương trình Coping Cat của Kendall (1994), được thiết kế cho trẻ 9–13 tuổi, có cấu trúc 16 buổi chia làm hai pha: 8 buổi đầu tập trung vào kỹ năng (nhận diện cảm xúc, kỹ thuật thư giãn, tái cấu trúc nhận thức, lập kế hoạch ứng phó), 8 buổi sau tập trung vào thực hành – tiếp xúc dần với các tình huống gây lo âu. Coping Cat đã được dịch và bản địa hóa ở hơn 15 quốc gia. Chương trình can thiệp gia đình của Barrett, Dadds và Rapee (1996) – tiền thân của giao thức FRIENDS – bổ sung thành phần làm việc với cha mẹ vào giao thức CBT cá nhân, dựa trên giả thuyết rằng cách phản ứng của cha mẹ trước các tình huống lo âu của con có thể duy trì hoặc làm trầm trọng triệu chứng. Nghiên cứu báo cáo nhóm có gia đình tham gia đạt tỷ lệ không còn chẩn đoán cao hơn nhóm CBT đơn thuần (84% so với 57%) sau 12 tháng theo dõi. Hai mô hình này đặt nền tảng cho hầu hết các giao thức CBT cho rối loạn lo âu trẻ em – vị thành niên hiện nay.'),
    ('P', 'Tại Việt Nam, các nghiên cứu can thiệp dựa trên CBT trên học sinh trung học cơ sở còn rất hạn chế. Luận án tiến sĩ y học của Trần Nguyễn Ngọc (2018) tại Trường Đại học Y Hà Nội đánh giá hiệu quả của liệu pháp thư giãn luyện tập — một thành phần hành vi của CBT — trên bệnh nhân rối loạn lo âu lan tỏa; mặc dù mẫu là người trưởng thành chứ không phải học sinh trung học cơ sở, công trình cung cấp dữ liệu về hiệu quả của kỹ thuật thư giãn trong bối cảnh Việt Nam. Chương trình Happy House do Tran và cộng sự (2023) thử nghiệm trong khuôn khổ hợp tác Việt – Úc đã đánh giá một can thiệp nâng cao sức khỏe tâm thần phổ quát theo mô hình thử nghiệm hai nhánh có đối chứng trên 1.084 học sinh trung học phổ thông tại Hà Nội. Tuy nhiên, các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu học sinh trung học cơ sở Việt Nam vẫn chưa được triển khai.'),
    ('P', 'Một số rào cản triển khai CBT tại Việt Nam cần được nhận diện: (i) khan hiếm chuyên gia tâm lý lâm sàng được đào tạo bài bản về CBT, đặc biệt cho trẻ em – vị thành niên; (ii) chi phí can thiệp CBT cá nhân 12–16 buổi vượt khả năng chi trả của phần lớn gia đình ở khu vực nông thôn; (iii) thiếu các bảng kiểm năng lực CBT được dịch và thích nghi cho ngữ cảnh Việt Nam, dẫn tới khó kiểm soát chất lượng triển khai khi mở rộng quy mô; (iv) thiếu các thang đo lo âu cho trẻ em được chuẩn hóa đầy đủ ở Việt Nam — phần lớn nghiên cứu hiện nay dùng các phiên bản tiếng Việt chưa qua quy trình thẩm định toàn diện.'),
    ('P', 'Bối cảnh văn hóa Việt Nam cũng đặt ra một số nội dung cần lưu ý khi bản địa hóa CBT. Thứ nhất là vai trò của gia đình mở rộng – ông bà, cô dì chú bác – trong việc nuôi dạy con cái; điều này khác với mô hình hạt nhân được giả định trong nhiều giao thức gốc của phương Tây. Việc huy động các thành viên gia đình mở rộng trong vai trò đồng minh có thể tăng độ tuân thủ giao thức, đặc biệt với các bài tập về nhà giữa các buổi. Thứ hai là khái niệm "bệnh tâm thần" trong nhận thức cộng đồng vẫn còn nặng tính kỳ thị, khiến quá trình tham gia điều trị thường bị chậm trễ và tâm lý tự đổ lỗi gia tăng. Thứ ba là tôn giáo và tín ngưỡng có thể đóng vai trò vừa bảo vệ vừa cản trở: niềm tin tâm linh có thể tạo nguồn an ủi, song cũng có thể dẫn tới việc tìm đến các biện pháp ngoài y học trước khi xem xét tham vấn lâm sàng (Praptomojati và cộng sự, 2024). Một giao thức CBT thích nghi văn hóa cho HSTHCS Việt Nam cần khai thác mặt tích cực của các yếu tố này – ví dụ huy động vai trò của ông bà như đồng minh trong việc duy trì bài tập về nhà – đồng thời chủ động giảm nhẹ các rào cản kỳ thị thông qua giáo dục cộng đồng đi kèm.'),

    ('H2', '4.2. CBT triển khai tại trường học'),
    ('P', 'Triển khai CBT tại trường mang lại nhiều ưu thế: tiếp cận số lượng lớn học sinh, giảm rào cản kỳ thị và chi phí, và cho phép can thiệp phổ quát ngay từ khi triệu chứng còn ở mức cận lâm sàng. Theo khung Mrazek-Haggerty (1994), kênh trường học là kênh duy nhất có thể đồng thời cung cấp can thiệp ở cả ba tầng phổ quát, chọn lọc và chỉ định trong cùng một không gian vật lý, qua đó tận dụng được hiệu ứng quy mô. Tuy nhiên, các thử nghiệm tại trường thường gặp khó khăn về độ trung thành với giao thức, khả năng đào tạo giáo viên và đo lường kết cục dài hạn.'),
    ('P', 'Phân tích tổng hợp của Cai và cộng sự (2025) trên 38 thử nghiệm lâm sàng ngẫu nhiên (21 thử nghiệm trong phân tích định lượng) về các chương trình tăng cường năng lực phục hồi tâm lý tại trường báo cáo hiệu ứng tổng hợp trên năng lực phục hồi đạt hiệu số trung bình chuẩn hóa (SMD) = 0,17 (khoảng tin cậy 95% từ 0,06 đến 0,29). Khi phân nhóm, các chương trình dựa trên chánh niệm cho hiệu ứng cao nhất (SMD = 0,57), tiếp đến là các chương trình lấy hoạt động thể thao làm trung tâm (SMD = 0,41). Mặc dù hiệu ứng tổng hợp nhỏ, các tác giả nhấn mạnh giá trị của can thiệp phổ quát quy mô lớn không qua sàng lọc – đặc biệt khi được tích hợp vào hoạt động trải nghiệm hoặc môn Giáo dục công dân; năng lực phục hồi được củng cố là một cấu trúc bảo vệ có liên quan đến việc giảm nguy cơ rối loạn lo âu.'),
    ('P', 'Thử nghiệm lâm sàng ngẫu nhiên cụm của Bradshaw và cộng sự (2025) đánh giá chương trình Early Adolescent Coping Power trên 709 học sinh lớp 7 tại 40 trường trung học cơ sở của Hoa Kỳ. Chương trình rút gọn gồm 25 buổi với học sinh và 12 buổi với cha mẹ, tích hợp các thành phần CBT cốt lõi. Theo dõi đến cuối lớp 8 — khoảng một năm sau can thiệp — cho thấy nhóm can thiệp giảm các vấn đề ngoại hiện do giáo viên đánh giá so với nhóm chứng, đồng thời ghi nhận cải thiện có ý nghĩa ở một số kết cục đối với học sinh nữ. Mặc dù đối tượng gốc của chương trình này là học sinh có hành vi gây hấn chứ không phải lo âu thuần, các thành phần CBT cốt lõi có thể được bản địa hóa để hướng tới lo âu trong bối cảnh Việt Nam.'),
    ('P', 'Tại Nhật Bản, Urao và cộng sự (2018, 2022) phát triển chương trình Journey of the Brave dựa trên CBT — phiên bản gốc 10 buổi × 45 phút (Urao và cộng sự, 2018) sau đó được mở rộng thành 14 buổi × 20 phút hoạt động ngắn trong lớp học (Urao và cộng sự, 2022) — thiết kế cho học sinh cuối tiểu học và được kiểm định qua các nghiên cứu tựa thực nghiệm có nhóm chứng, ghi nhận hiệu ứng giảm triệu chứng lo âu. Việc rút ngắn mỗi buổi từ 45 phút xuống 20 phút trong bản 2022 là một thay đổi thiết kế đáng chú ý — tận dụng được khe giờ ngắn giữa các tiết học chính khóa, qua đó giảm gánh nặng cho thời gian biểu của trường. Mô hình "buổi ngắn lồng ghép" này có giá trị tham khảo cao cho thiết kế chương trình tại Việt Nam, nơi quỹ thời gian dành cho hoạt động tâm lý ngoài giờ học rất hạn chế.'),
    ('P', 'Tại Trung Quốc, He và cộng sự (2026) báo cáo một nghiên cứu thử nghiệm chương trình CBT trường học ngắn (Power Up – CBTD) cho học sinh có triệu chứng trầm cảm, công bố trên tạp chí Journal of Affective Disorders; chương trình cải thiện triệu chứng trầm cảm, trong khi triệu chứng lo âu có xu hướng giảm nhưng khác biệt giữa hai nhóm chưa đạt ý nghĩa thống kê — cho thấy hiệu quả lan tỏa sang lo âu của các can thiệp tập trung vào trầm cảm vẫn cần được kiểm chứng thêm. Mẫu nghiên cứu gồm học sinh lớp 5, 6 và 9 tại tỉnh Hà Nam (Trung Quốc) — một thiết kế đa lứa tuổi thú vị nhưng cũng đặt ra câu hỏi về tính đồng nhất của hiệu ứng giữa các nhóm tuổi.'),
    ('P', 'Tại Việt Nam, các chương trình CBT tại trường còn ở giai đoạn rất sớm. Phần lớn các trường THCS hiện nay đã có tổ tư vấn tâm lý theo Thông tư 31/2017/TT-BGDĐT, nhưng hoạt động chủ yếu là tham vấn cá nhân khi học sinh chủ động tìm đến, không phải chương trình phổ quát có cấu trúc. Khoảng trống lớn nhất tại Việt Nam là sự thiếu vắng một chương trình CBT bản địa được kiểm định trên học sinh trung học cơ sở, thiết kế phù hợp với khung chương trình môn Hoạt động trải nghiệm – Hướng nghiệp của Chương trình giáo dục phổ thông 2018. Một mô hình hai tầng có thể tham khảo: tầng phổ quát (4–6 buổi tích hợp môn học chính khóa, tập trung vào kỹ năng nhận diện cảm xúc) và tầng có chọn lọc (8–12 buổi cho học sinh được sàng lọc có triệu chứng cận lâm sàng). Việc huy động sự đồng hành của phụ huynh trong quá trình can thiệp cũng là một yếu tố cần được cân nhắc khi thiết kế chương trình cho bối cảnh văn hóa Việt Nam.'),
    ('P', 'Vai trò của giáo viên chủ nhiệm trong triển khai CBT tại trường cũng cần được làm rõ. Khác với chuyên gia tâm lý lâm sàng, giáo viên chủ nhiệm có lợi thế về tần suất tiếp xúc với học sinh và hiểu biết về hoàn cảnh gia đình, song lại hạn chế về thời gian và năng lực can thiệp lâm sàng. Mô hình "giáo viên chủ nhiệm là điều phối viên – chuyên gia tâm lý là người giám sát" đã được áp dụng tại một số nước châu Á có thể là phương án phù hợp cho điều kiện Việt Nam khi nguồn lực chuyên môn còn hạn chế. Cụ thể, giáo viên chủ nhiệm có thể được đào tạo về sàng lọc sơ cấp và kỹ thuật can thiệp cơ bản trong khuôn khổ chương trình môn học, trong khi nhà tâm lý đảm nhận giám sát chuyên môn, tư vấn ca khó và tiếp nhận các trường hợp cần can thiệp sâu hơn. Việc phân tầng vai trò như vậy giúp tận dụng tối đa nguồn lực sẵn có, đồng thời bảo đảm chất lượng can thiệp cho các trường hợp phức tạp.'),
    ('P', 'Một yếu tố thiết kế quan trọng cho chương trình CBT trường học là tích hợp thành phần làm việc với cha mẹ. Mô hình của Barrett, Dadds và Rapee (1996) cho thấy việc thêm các buổi làm việc với gia đình cho hiệu quả cải thiện đáng kể so với CBT cá nhân đơn thuần, đặc biệt ở lứa tuổi đầu vị thành niên khi học sinh vẫn phụ thuộc nhiều vào gia đình. Trong bối cảnh Việt Nam — nơi gia đình mở rộng có ảnh hưởng lớn tới đời sống tinh thần học sinh — thành phần này có thể có giá trị tăng thêm. Tuy nhiên, việc huy động cha mẹ tham gia các buổi tại trường gặp khó khăn do lịch làm việc, đặc biệt ở khu vực nông thôn nơi nhiều phụ huynh đi làm xa. Một giải pháp thay thế khả thi là cung cấp tài liệu cho phụ huynh qua các nhóm Zalo/Facebook của lớp, kết hợp 1–2 buổi gặp trực tiếp được tổ chức vào cuối tuần.'),
    ('P', 'Đánh giá tính bền vững của các chương trình CBT trường học sau khi giai đoạn nghiên cứu kết thúc cũng là một thách thức quan trọng. Theo các tổng quan quốc tế, chỉ khoảng 30–40% các chương trình can thiệp được duy trì sau khi nguồn tài trợ nghiên cứu chấm dứt; phần còn lại bị thu hẹp hoặc ngừng hẳn. Các yếu tố thúc đẩy bền vững gồm: (i) sự ủng hộ của lãnh đạo trường; (ii) tích hợp chương trình vào khung môn học chính thức thay vì hoạt động ngoại khóa; (iii) cơ chế đào tạo liên tục cho giáo viên mới; (iv) nguồn tài chính ổn định từ ngân sách nhà nước hoặc các đối tác dài hạn. Khi thiết kế các thử nghiệm CBT trường học tại Việt Nam, các yếu tố này cần được lên kế hoạch ngay từ giai đoạn đầu chứ không chờ tới khi giai đoạn nghiên cứu kết thúc.'),

    ('H2', '4.3. CBT qua nền tảng số'),
    ('P', 'Can thiệp CBT qua nền tảng số đang phát triển nhanh nhất trong lĩnh vực can thiệp lo âu vị thành niên. DMHI bao quát nhiều hình thức – từ ứng dụng độc lập, chương trình internet có hướng dẫn, đến nền tảng kết hợp giữa gặp mặt trực tiếp và trực tuyến (blended). Ưu thế lớn nhất của kênh số là khả năng tiếp cận đa số học sinh trong khi cần ít nhân lực chuyên môn; nhược điểm chính là độ tuân thủ thấp và đòi hỏi cơ sở hạ tầng công nghệ.'),
    ('P', 'Phân tích tổng hợp của Walder và cộng sự (2025) trên 21 thử nghiệm lâm sàng ngẫu nhiên về can thiệp số cho rối loạn lo âu xã hội ở trẻ em, vị thành niên và thanh niên trẻ báo cáo hiệu ứng tổng hợp g = 0,508 (khoảng tin cậy 95% từ 0,31 đến 0,71). Khi phân nhóm theo mức độ hỗ trợ, các can thiệp số có hướng dẫn đạt hiệu ứng g = 0,825, cao hơn rõ rệt so với các can thiệp không hướng dẫn (g = 0,27) – cho thấy thành phần hỗ trợ chuyên môn, dù tối thiểu, vẫn đóng vai trò quan trọng đối với hiệu quả của can thiệp số. Phát hiện này có hai hàm ý quan trọng: thứ nhất, các thiết kế DMHI hoàn toàn tự động không phải lúc nào cũng tối ưu — một mức độ tương tác con người tối thiểu giúp nâng cao hiệu quả đáng kể; thứ hai, mô hình blended (tương tác trực tuyến + buổi gặp định kỳ) có thể là điểm cân bằng tốt giữa quy mô và hiệu quả.'),
    ('P', 'Matsumoto và cộng sự (2024) đánh giá một chương trình CBT qua internet không hướng dẫn trên điện thoại cho vị thành niên và thanh niên trẻ Nhật Bản có rối loạn lo âu xã hội dưới ngưỡng, công bố trên tạp chí JMIR Pediatrics and Parenting. Thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu 77 người tham gia 15–25 tuổi (38 can thiệp + 39 đối chứng) báo cáo hiệu ứng g = 0,66 đối với điểm lo âu xã hội sau 10 tuần can thiệp – con số đáng chú ý vì can thiệp gần như không có sự tham gia trực tiếp của chuyên gia. Yếu tố thiết kế tương tác — bao gồm các bài tập tương tác ngắn, phản hồi cá nhân hóa và lịch nhắc — được xác định là chìa khóa duy trì sự tham gia của người dùng.'),
    ('P', 'Thử nghiệm lâm sàng ngẫu nhiên của Bress và cộng sự (2024) công bố trên JAMA Network Open đánh giá ứng dụng Maya – một nền tảng CBT tự hướng dẫn 12 phiên trên điện thoại – trên 59 thanh niên 18–25 tuổi với rối loạn lo âu (rối loạn lo âu tổng quát 56%, rối loạn lo âu xã hội 41%). Kết quả ghi nhận hiệu ứng kích thước rất lớn trên các thang đo phụ (Anxiety Sensitivity Index Cohen d = 0,93 tại tuần 6; Liebowitz Social Anxiety Scale Cohen d = 1,07 tại tuần 12). Mặc dù mẫu chính là thanh niên trẻ chứ không phải HSTHCS, kết quả này gợi ý CBT số được thiết kế chu đáo có thể đạt hiệu quả lớn và cần được điều chỉnh cho lứa tuổi nhỏ hơn. Một hướng thiết kế đáng tham khảo là ứng dụng ClearlyMe — một nền tảng CBT số được đồng thiết kế cùng vị thành niên cho trầm cảm và lo âu (Li và cộng sự, 2022), do Black Dog Institute thuộc Đại học New South Wales (Úc) phát triển.'),
    ('P', 'Tại Việt Nam, chưa có nền tảng DMHI tiếng Việt nào được kiểm định bằng thử nghiệm lâm sàng ngẫu nhiên trên HSTHCS. Một số nỗ lực bao gồm các nền tảng tham vấn trực tuyến do các tổ chức phi chính phủ vận hành, song các nền tảng này thường thiếu thành phần CBT có cấu trúc. Trong khi đó, tỷ lệ HSTHCS Việt Nam sở hữu điện thoại thông minh đã đạt mức bão hòa ở khu vực đô thị và đang tăng nhanh ở khu vực nông thôn, tạo điều kiện thuận lợi cho việc triển khai DMHI một khi nội dung phù hợp được phát triển.'),
    ('P', 'Các nghiên cứu cần lưu ý một số rủi ro đặc thù của DMHI: (i) tỷ lệ học sinh hoàn thành đầy đủ giao thức trong các nghiên cứu quốc tế thường chỉ đạt 40–60%; (ii) bảo mật dữ liệu cá nhân, đặc biệt với vị thành niên; (iii) bất bình đẳng số – học sinh thiếu kết nối internet hoặc gia đình không cho phép sử dụng điện thoại sẽ bị loại trừ; (iv) DMHI không thể thay thế cho can thiệp lâm sàng có chuyên gia khi học sinh có triệu chứng nặng hoặc nguy cơ tự sát; cần xây dựng quy trình chuyển tiếp rõ ràng từ DMHI sang dịch vụ chuyên khoa.'),
    ('P', 'Một thảo luận hữu ích cho Việt Nam là so sánh ba mô hình triển khai DMHI tiềm năng. Mô hình thứ nhất là ứng dụng độc lập đặt trên kho ứng dụng cho học sinh tự tải, với chi phí thấp nhưng độ duy trì hạn chế. Mô hình thứ hai là ứng dụng tích hợp vào nền tảng học tập đã có như VNEDU hoặc các hệ thống quản lý học tập của nhà trường, qua đó tận dụng cơ chế đăng nhập và thông báo đã sẵn có. Mô hình thứ ba là ứng dụng kết hợp gặp mặt – trực tuyến (blended) với điều phối viên là giáo viên tâm lý học đường, cân bằng giữa độ tiếp cận và chất lượng chuyên môn. Mỗi mô hình có ưu nhược điểm và phù hợp với từng phân khúc HSTHCS khác nhau — khu vực đô thị có nguồn lực tốt có thể triển khai mô hình thứ ba, trong khi khu vực nông thôn và miền núi mô hình thứ nhất có thể là điểm khởi đầu khả thi.'),
    ('P', 'Khi mở rộng quy mô triển khai DMHI tại Việt Nam, một số khía cạnh kỹ thuật cần được lập kế hoạch trước. Thứ nhất là cơ chế bảo mật dữ liệu: dữ liệu sức khỏe tâm thần của vị thành niên là loại thông tin nhạy cảm, đòi hỏi tuân thủ các nguyên tắc tối thiểu hóa thu thập, mã hóa khi lưu trữ, và minh bạch về thời gian giữ dữ liệu. Thứ hai là cơ chế cảnh báo khi học sinh có dấu hiệu tự sát hoặc tự hại: nền tảng cần có quy trình tự động phát hiện và chuyển tiếp tới cán bộ tâm lý hoặc cha mẹ – một yếu tố không thể thay bằng thuật toán đơn giản. Thứ ba là cơ chế đánh giá hiệu quả liên tục: cần thu thập dữ liệu sử dụng (engagement metrics) và dữ liệu kết cục theo thời gian thực để cho phép cải tiến lặp đi lặp lại của nền tảng. Các yêu cầu này đòi hỏi sự phối hợp chuyên môn giữa nhà tâm lý lâm sàng, kỹ sư phần mềm và chuyên gia bảo mật ngay từ giai đoạn thiết kế.'),
    ('P', 'Một bài học quan trọng từ các thử nghiệm DMHI quốc tế là vai trò của giai đoạn đồng thiết kế (co-design) với người dùng cuối. Li và cộng sự (2022) khi xây dựng ClearlyMe đã tổ chức nhiều vòng góp ý với vị thành niên Úc trong giai đoạn phát triển – một quy trình kéo dài hơn một năm trước khi ứng dụng được phát hành công khai. Quá trình đồng thiết kế giúp các thành phần CBT vốn có nguồn gốc lâm sàng được "dịch" sang ngôn ngữ, hình ảnh và bối cảnh tương tác phù hợp với vị thành niên. Đối với Việt Nam, một dự án phát triển nền tảng DMHI tiếng Việt cần dành ít nhất 6–9 tháng cho giai đoạn đồng thiết kế trước khi bắt đầu thử nghiệm chính thức, với sự tham gia của các nhóm học sinh thuộc các vùng địa lý và bối cảnh kinh tế – xã hội khác nhau.'),
    ('P', 'Bằng chứng so sánh giữa hai dạng can thiệp số quan trọng — tự hướng dẫn (self-guided) và có hướng dẫn (guided) — tiếp tục được củng cố. Walder và cộng sự (2025) ghi nhận sự chênh lệch lớn giữa hai dạng (g = 0,825 so với g = 0,27), trong khi Matsumoto và cộng sự (2024) cho thấy một thiết kế tự hướng dẫn tốt vẫn có thể đạt g = 0,66 nếu các thành phần tương tác được thiết kế khéo léo. Hai phát hiện này không mâu thuẫn — chúng cho thấy chất lượng thiết kế là yếu tố quyết định: một ứng dụng tự hướng dẫn tốt có thể tiệm cận một ứng dụng có hướng dẫn yếu, nhưng nói chung sự hiện diện của một mức độ hỗ trợ chuyên môn tối thiểu vẫn là tài sản đáng giá. Đối với Việt Nam, mô hình can thiệp số có hướng dẫn từ xa (qua nền tảng chat hoặc video call định kỳ) bởi sinh viên thạc sĩ tâm lý lâm sàng có thể là phương án vừa khả thi vừa hiệu quả.'),

    ('H2', '4.4. Mô hình kết hợp và chăm sóc theo bậc (stepped care)'),
    ('P', 'Ngoài việc triển khai từng kênh riêng lẻ, các mô hình kết hợp giữa ba kênh đang được khuyến khích để tối ưu hiệu quả và chi phí. Mô hình chăm sóc theo bậc (stepped care) của Bower và Gilbody (2005), ban đầu được phát triển cho hệ thống dịch vụ Tâm lý điều trị tăng cường (Improving Access to Psychological Therapies – IAPT) của Vương quốc Anh, được tổ chức theo nguyên tắc: bắt đầu bằng can thiệp ít cường độ nhất phù hợp với mức triệu chứng, leo bậc lên can thiệp cường độ cao hơn khi không đáp ứng. Áp dụng vào bối cảnh lo âu vị thành niên, mô hình stepped care có thể vận hành như sau: bậc 1 là tài liệu tự giúp hoặc ứng dụng số không hướng dẫn; bậc 2 là CBT nhóm tại trường hoặc can thiệp số có hướng dẫn; bậc 3 là CBT cá nhân với chuyên gia. Các nghiên cứu cho thấy mô hình này có hiệu quả tương đương với mô hình chăm sóc truyền thống nhưng tiết kiệm khoảng 20–30% chi phí nguồn lực chuyên môn.'),
    ('P', 'Mô hình can thiệp pha trộn (blended intervention) là dạng kết hợp thứ hai đáng chú ý. Trong blended, học sinh tham gia một phần qua nền tảng số (đọc tài liệu, làm bài tập tương tác, theo dõi tâm trạng) và một phần qua các buổi gặp trực tiếp định kỳ với giáo viên tâm lý hoặc chuyên gia (1–2 lần/tháng). Mô hình này có ưu thế ở việc duy trì độ tuân thủ thông qua các buổi gặp đồng thời tận dụng tính linh hoạt của nền tảng số. Các thử nghiệm sơ bộ cho thấy mô hình blended có hiệu quả tương đương CBT cá nhân truyền thống nhưng giảm 40–60% tổng thời gian tiếp xúc chuyên gia.'),
    ('P', 'Đối với Việt Nam, mô hình stepped care kết hợp blended có thể là phương án triển khai khả thi nhất trong giai đoạn 2026–2030. Cụ thể, ở cấp trường THCS có thể vận hành ba bậc: bậc 1 (phổ quát) là tài liệu tâm lý lồng ghép vào môn Hoạt động trải nghiệm – Hướng nghiệp + ứng dụng số miễn phí; bậc 2 (chọn lọc/chỉ định nhẹ) là CBT nhóm 6–8 buổi do giáo viên tâm lý dẫn dắt, kết hợp ứng dụng số làm bài tập về nhà; bậc 3 (chỉ định nặng) là CBT cá nhân tại trung tâm tâm lý tuyến tỉnh hoặc bệnh viện chuyên khoa. Mô hình này cho phép sử dụng hiệu quả nguồn nhân lực chuyên môn vốn rất hạn chế của Việt Nam, đồng thời tạo cơ chế chuyển tuyến rõ ràng giữa các tầng.'),
    ('P', 'Một thách thức quan trọng đối với mô hình stepped care là việc xác định ngưỡng "không đáp ứng" để chuyển bậc. Trong các hệ thống quốc tế, ngưỡng này thường được đặt theo tỷ lệ giảm điểm trên các thang đo chuẩn hóa sau một thời gian can thiệp (ví dụ giảm dưới 50% sau 6 tuần). Tại Việt Nam, chưa có quy trình chuẩn hóa nào cho ngưỡng này, và việc xây dựng quy trình đó là một trong những nhiệm vụ cần ưu tiên khi triển khai mô hình stepped care.'),
    ('CAP', 'Bảng 2. So sánh ba kênh can thiệp CBT cho HSTHCS'),
    ('TABLE', ['Tiêu chí', 'Kênh lâm sàng', 'Kênh trường học', 'Kênh số'],
              [['Hiệu quả (cỡ effect tổng hợp)', 'd = 0,80–1,07 (Bress 2024 spec.)', 'SMD = 0,17–0,57 (Cai 2025)', 'g = 0,27–0,83 (Walder 2025)'],
               ['Khả năng tiếp cận', 'Thấp', 'Cao', 'Rất cao'],
               ['Chi phí cho mỗi học sinh', 'Cao', 'Thấp', 'Rất thấp'],
               ['Yêu cầu nhân lực chuyên môn', 'Rất cao', 'Vừa', 'Thấp (nếu tự hướng dẫn)'],
               ['Độ tuân thủ', 'Cao (~80%)', 'Vừa (~60–70%)', 'Thấp (~40–60%)'],
               ['Phù hợp tầng can thiệp', 'Chỉ định, lâm sàng', 'Phổ quát, chọn lọc, chỉ định nhẹ', 'Phổ quát, chọn lọc']],
              [3.5, 4.0, 4.0, 4.0]),
]

SEC_BAN_LUAN = [
    ('H1', '5. Bàn luận'),

    ('H2', '5.1. Tổng hợp ba nhóm can thiệp'),
    ('P', 'Tổng hợp các bằng chứng đã trình bày, ba nhóm can thiệp CBT cho rối loạn lo âu ở HSTHCS hiện có nền tảng bằng chứng quốc tế khá vững. CBT lâm sàng có cỡ hiệu ứng lớn nhất nhưng khả năng tiếp cận thấp. CBT trường học có cỡ hiệu ứng tổng hợp nhỏ tới vừa nhưng phù hợp với mô hình can thiệp phổ quát và chọn lọc. CBT số có cỡ hiệu ứng vừa và đang phát triển nhanh, đặc biệt khi kết hợp một mức độ hỗ trợ chuyên môn tối thiểu. Tại Việt Nam, hầu hết các báo cáo còn ở dạng thử nghiệm pilot quy mô nhỏ hoặc nghiên cứu lý luận chưa được củng cố bằng dữ liệu thực nghiệm đối chứng; đây là khoảng cách lớn so với các nước có hệ thống chăm sóc sức khỏe tâm thần phát triển.'),
    ('P', 'Bức tranh phân tầng theo khung Mrazek-Haggerty (1994) cho thấy điểm yếu rõ rệt nhất của hệ thống can thiệp Việt Nam hiện nay nằm ở tầng phổ quát và chọn lọc — hai tầng vốn cần được tổ chức qua trường học. Trong khi tầng chỉ định (can thiệp lâm sàng) có một số nỗ lực qua các bệnh viện chuyên khoa và phòng tham vấn tâm lý tư nhân, hai tầng còn lại gần như chưa có chương trình có cấu trúc nào. Sự khuyết tật này đặc biệt nghiêm trọng vì hai tầng phổ quát và chọn lọc là nơi có thể can thiệp sớm khi triệu chứng còn ở mức cận lâm sàng, ngăn ngừa diễn tiến tới rối loạn lâm sàng đầy đủ — tức là nơi có hiệu suất chi phí cao nhất trong toàn hệ thống chăm sóc.'),
    ('P', 'Một điểm chung quan trọng giữa ba nhóm là vai trò của các thành phần CBT cốt lõi — tái cấu trúc nhận thức, tiếp xúc dần, kỹ thuật thư giãn — vẫn được giữ nguyên dù kênh cung cấp khác nhau. Điều này gợi ý rằng cốt lõi giao thức CBT là khá bền vững giữa các kênh, và sự khác biệt về hiệu quả phần lớn đến từ liều lượng (số buổi, độ dài mỗi buổi, có hỗ trợ chuyên gia hay không) và mức độ tương tác — chứ không phải từ thành phần lý thuyết. Phát hiện này quan trọng cho thiết kế tại Việt Nam: việc bản địa hóa nội dung văn hóa và ngôn ngữ phải đi đôi với việc duy trì các thành phần cốt lõi để bảo đảm hiệu quả lâm sàng.'),

    ('H2', '5.2. Năm khoảng trống chính trong nghiên cứu can thiệp tại Việt Nam'),
    ('P', 'Trên cơ sở khung phân tích ba kênh × ba tầng, có thể nhận diện năm khoảng trống chính trong nghiên cứu can thiệp cho HSTHCS Việt Nam.'),
    ('P', 'Thứ nhất, thiếu thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu HSTHCS — phần lớn các nghiên cứu hiện có đều dùng thiết kế trước-sau hoặc tựa thực nghiệm với nhóm chứng không phân ngẫu nhiên. Hệ quả là chưa có ước tính hiệu ứng đáng tin cậy cho riêng đối tượng và bối cảnh Việt Nam, dẫn tới các khuyến nghị chính sách thiếu cơ sở định lượng. Việc thiết kế và thực hiện ít nhất một thử nghiệm RCT đa trung tâm trong giai đoạn 2026–2030 là điều kiện cần để đặt nền tảng bằng chứng cho hệ thống can thiệp quốc gia.'),
    ('P', 'Thứ hai, thiếu chương trình CBT dựa vào trường được thiết kế phù hợp với chương trình giáo dục phổ thông Việt Nam — mặc dù chương trình giáo dục phổ thông 2018 đã đưa môn Hoạt động trải nghiệm – Hướng nghiệp vào chương trình chính thức, các giáo viên hiện chưa được đào tạo có hệ thống về sàng lọc và can thiệp sơ cấp ở mức cận lâm sàng. Một số trường có biên chế giáo viên tâm lý học đường nhưng năng lực thực sự về CBT cho lo âu trẻ em – vị thành niên vẫn còn hạn chế.'),
    ('P', 'Thứ ba, thiếu nền tảng CBT số tiếng Việt được kiểm định — các nền tảng tiếng Anh hiện có như Maya hay ClearlyMe không thể triển khai trực tiếp do rào cản ngôn ngữ và văn hóa. Phát triển một nền tảng CBT số tiếng Việt đòi hỏi sự phối hợp giữa nhà tâm lý, kỹ sư phần mềm và chuyên gia thiết kế trải nghiệm người dùng — một mô hình hợp tác đa ngành chưa phổ biến trong cộng đồng nghiên cứu sức khỏe tâm thần Việt Nam.'),
    ('P', 'Thứ tư, thiếu các nghiên cứu về thiết kế và đo lường kết cục dài hạn. Các nghiên cứu hiện có chủ yếu báo cáo kết cục ngay sau can thiệp, ít có theo dõi 6 tháng, 12 tháng hoặc dài hơn để đánh giá tính bền vững. Đối với rối loạn lo âu — vốn có xu hướng tái phát — việc thiếu dữ liệu theo dõi dài hạn là khoảng trống nghiêm trọng.'),
    ('P', 'Thứ năm, thiếu công cụ sàng lọc và đo lường chuẩn hóa tiếng Việt cho rối loạn lo âu vị thành niên. Mặc dù một số thang đo như RCADS, SCARED, DASS-Y đã có phiên bản tiếng Việt, quy trình thẩm định tâm trắc toàn diện (độ tin cậy, độ hiệu lực hội tụ, độ hiệu lực phân biệt, ngưỡng cắt cho dân số Việt Nam) chưa được hoàn tất cho phần lớn các công cụ này. Việc thiếu công cụ chuẩn hóa làm cho các nghiên cứu can thiệp khó so sánh kết quả với nhau và với các nghiên cứu quốc tế.'),

    ('H2', '5.3. Khác biệt vùng miền và yếu tố giới tính trong can thiệp'),
    ('P', 'Sự khác biệt vùng miền là một yếu tố cấu trúc cần lưu ý khi triển khai can thiệp tại Việt Nam. Khu vực đô thị có thể triển khai đầy đủ cả ba tầng can thiệp (lâm sàng, trường học, số) với nguồn lực chuyên môn tốt; khu vực nông thôn và miền núi có thể chỉ tiếp cận được phần can thiệp tại trường và can thiệp số đơn giản. Vì vậy, các chiến lược can thiệp tương lai nên được phân tầng theo đặc thù vùng miền, tránh áp dụng đồng nhất một mô hình duy nhất. Đối với khu vực có nhiều học sinh dân tộc thiểu số, cần đặc biệt cân nhắc yếu tố ngôn ngữ và văn hóa khi chuyển giao chương trình – chẳng hạn dịch tài liệu sang tiếng dân tộc khi cần, đào tạo giáo viên người dân tộc làm điều phối viên và xây dựng vai trò của già làng – người có uy tín cộng đồng trong việc giảm kỳ thị.'),
    ('P', 'Cũng cần lưu ý sự khác biệt giữa các vùng miền không chỉ về kinh tế – cơ sở hạ tầng mà còn về mẫu hình áp lực mà học sinh phải đối mặt. Học sinh khu vực đô thị thường chịu áp lực học tập từ kỳ vọng cao của gia đình, áp lực thành tích để vào trường chất lượng cao; trong khi học sinh khu vực nông thôn có thể chịu các áp lực khác như cha mẹ đi làm xa, thiếu hỗ trợ tâm lý từ gia đình, tiếp xúc thiếu cấu trúc với internet và mạng xã hội. Hai mẫu hình áp lực này dẫn tới các biểu hiện lo âu khác nhau, đòi hỏi các phiên bản giao thức can thiệp được điều chỉnh phù hợp. Một thiết kế nghiên cứu can thiệp có giá trị cao là so sánh hiệu quả của cùng một giao thức CBT khi triển khai trên hai nhóm dân số đối lập này, qua đó định lượng được mức độ bản địa hóa cần thiết.'),
    ('P', 'Về yếu tố giới tính, các nghiên cứu quốc tế nhất quán cho thấy nữ giới có tỷ lệ rối loạn lo âu cao hơn nam giới khoảng 1,5–2 lần ở lứa tuổi vị thành niên, đồng thời nữ cũng có xu hướng đáp ứng tốt hơn với can thiệp CBT nói chung. Bradshaw và cộng sự (2025) đã ghi nhận sự khác biệt theo giới trong đáp ứng can thiệp Coping Power, với học sinh nữ cải thiện rõ rệt hơn ở một số kết cục. Tại Việt Nam, dữ liệu phân tích theo giới còn rất hạn chế; các thiết kế nghiên cứu tương lai cần tích hợp phân tích phân nhóm theo giới ngay từ giai đoạn lập kế hoạch.'),

    ('H2', '5.4. Khoảng trống về thiết kế nghiên cứu can thiệp'),
    ('P', 'Đa số các nghiên cứu can thiệp tại Việt Nam đang sử dụng thiết kế trước-sau (pre-post) một nhóm hoặc thiết kế tựa thực nghiệm với nhóm chứng không phân ngẫu nhiên. Các thiết kế này không cho phép quy kết nhân quả chắc chắn cho hiệu ứng quan sát được, do không loại trừ được các yếu tố nhiễu như sự cải thiện tự nhiên theo thời gian, hiệu ứng kỳ vọng và hiệu ứng đo lường lặp lại. Để nâng cao chất lượng bằng chứng, các thiết kế nghiên cứu tương lai cần áp dụng: (i) phân nhóm ngẫu nhiên (random allocation); (ii) đối chứng tích cực thay vì chỉ đối chứng đợi; (iii) phép đo mù (blind assessment) khi đánh giá kết cục; (iv) đăng ký trước thiết kế nghiên cứu trên các nền tảng như ClinicalTrials.gov hoặc Open Science Framework để tăng tính minh bạch.'),
    ('P', 'Thiết kế thử nghiệm lâm sàng ngẫu nhiên cụm (cluster RCT) – như mô hình Bradshaw và cộng sự (2025) – đặc biệt phù hợp cho can thiệp tại trường, vì việc phân ngẫu nhiên ở cấp trường tránh được hiệu ứng lây lan giữa các học sinh cùng lớp. Đây là thiết kế nên được ưu tiên cho các thử nghiệm CBT trường học tại Việt Nam trong tương lai.'),

    ('H2', '5.5. Khoảng trống về đo lường hiệu quả và kết cục dài hạn'),
    ('P', 'Bốn vấn đề về đo lường cần được giải quyết khi xây dựng các nghiên cứu can thiệp tại Việt Nam. Thứ nhất, lựa chọn công cụ đo lường: cần ưu tiên các thang đo đã được thẩm định tâm trắc đầy đủ ở Việt Nam (RCADS, SCARED, DASS-Y phiên bản tiếng Việt) thay vì các thang đo tự tạo chưa có dữ liệu độ tin cậy. Thứ hai, thời điểm đo lường: ngoài đo trước–sau can thiệp, cần thêm các thời điểm theo dõi 3 tháng, 6 tháng và 12 tháng để đánh giá độ bền vững. Thứ ba, đa nguồn đánh giá: kết hợp tự báo cáo của học sinh, đánh giá của cha mẹ và giáo viên để giảm thiên kiến tự báo cáo. Thứ tư, các kết cục thứ cấp: ngoài điểm thang đo, cần đo các kết cục có ý nghĩa thực tế như mức độ tham gia hoạt động xã hội, chất lượng giấc ngủ, kết quả học tập.'),
    ('P', 'Đặc biệt, các nghiên cứu can thiệp số cần báo cáo các chỉ số tương tác (engagement metrics) như tỷ lệ hoàn thành chương trình, số ngày sử dụng trung bình, tỷ lệ học sinh hoàn thành bài tập về nhà — những chỉ số này quan trọng để hiểu cơ chế và để cải tiến nền tảng. Hiện nay rất ít nghiên cứu DMHI tại Việt Nam báo cáo các chỉ số này một cách hệ thống.'),

    ('H2', '5.6. Hàm ý chính sách và chuyển giao tri thức'),
    ('P', 'Trên cơ sở các khoảng trống đã nhận diện, ba hàm ý chính sách nổi lên. Thứ nhất, cần một cơ chế tài trợ nghiên cứu can thiệp dài hạn — không chỉ tài trợ cho thử nghiệm đơn lẻ mà tài trợ cho cả mạng lưới nhiều thử nghiệm song song trên các vùng miền khác nhau. Cơ chế này có thể được tổ chức qua Quỹ Phát triển Khoa học và Công nghệ Quốc gia (NAFOSTED) hoặc qua các chương trình mục tiêu quốc gia về sức khỏe.'),
    ('P', 'Thứ hai, cần một khung pháp lý rõ ràng cho việc tích hợp dịch vụ tâm lý vào hệ thống giáo dục. Thông tư 31/2017/TT-BGDĐT về tổ tư vấn tâm lý đã đặt nền tảng ban đầu, nhưng cần được bổ sung bằng các quy định chi tiết về (i) chuẩn năng lực của giáo viên tâm lý học đường; (ii) quy trình sàng lọc – chuyển tuyến giữa trường học và cơ sở y tế; (iii) cơ chế phối hợp với phụ huynh.'),
    ('P', 'Thứ ba, cần một cơ chế chuyển giao tri thức từ nghiên cứu sang thực hành. Hiện nay, khoảng cách giữa các nghiên cứu được công bố và việc áp dụng chúng vào hệ thống dịch vụ rất lớn. Việc thiết lập các mạng lưới chuyển giao — bao gồm các diễn đàn chuyên môn định kỳ, các khóa tập huấn quy chuẩn, và các tài liệu hướng dẫn thực hành dựa trên bằng chứng — là điều kiện cần để nghiên cứu thực sự có tác động lên đời sống học sinh.'),

    ('H2', '5.7. Năm hướng nghiên cứu ưu tiên cho giai đoạn 2026–2030'),
    ('P', 'Trên cơ sở các khoảng trống đã nhận diện, nhóm tác giả đề xuất năm hướng nghiên cứu ưu tiên cho giai đoạn 2026–2030.'),
    ('P', 'Hướng thứ nhất, ưu tiên thiết kế và thực hiện một thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu HSTHCS Việt Nam với giao thức CBT đã được bản địa hóa. Thử nghiệm cần có cỡ mẫu đủ lớn (ước tính ≥ 800 học sinh phân chia ở 8–12 trường) để cho phép phân tích phân nhóm theo giới và theo vùng miền, kèm theo dõi tối thiểu 12 tháng sau can thiệp.'),
    ('P', 'Hướng thứ hai, phát triển một chương trình CBT trường học 8–12 buổi tích hợp vào môn Hoạt động trải nghiệm – Hướng nghiệp, với tài liệu hướng dẫn dành cho giáo viên chủ nhiệm và giáo viên tâm lý học đường. Chương trình cần bao gồm cả phần tài liệu cho học sinh, sổ tay cho giáo viên, và bộ công cụ sàng lọc – đánh giá kết cục đi kèm.'),
    ('P', 'Hướng thứ ba, xây dựng một nền tảng CBT số tiếng Việt theo mô hình của các ứng dụng đã được phát triển trên thế giới như Maya (Bress và cộng sự, 2024) hoặc ClearlyMe (Li và cộng sự, 2022), kết hợp với mạng lưới chuyển tiếp tới dịch vụ chuyên khoa khi cần. Nền tảng cần đáp ứng các yêu cầu bảo mật dữ liệu, cơ chế cảnh báo tự sát, và đánh giá hiệu quả liên tục thông qua dữ liệu sử dụng.'),
    ('P', 'Hướng thứ tư, đào tạo bài bản cho đội ngũ giáo viên tâm lý học đường về sàng lọc lo âu, can thiệp sơ cấp và đánh giá đầu ra. Chương trình đào tạo nên có ba mức: chứng chỉ căn bản (40 giờ) cho giáo viên chủ nhiệm; chứng chỉ chuyên môn (200 giờ) cho giáo viên tâm lý học đường; và chứng chỉ chuyên sâu (sau đại học) cho chuyên gia tâm lý lâm sàng.'),
    ('P', 'Hướng thứ năm, thiết lập một mạng lưới nghiên cứu can thiệp Việt Nam kết nối các trường đại học, viện nghiên cứu và bệnh viện tâm thần để chuẩn hóa thang đo, giao thức và quy trình báo cáo. Mạng lưới này có thể tổ chức theo mô hình consortium với hội đồng điều phối quay vòng giữa các đơn vị thành viên.'),

    ('H2', '5.8. Phối hợp ba cấp độ để các hướng nghiên cứu khả thi'),
    ('P', 'Để các khuyến nghị này khả thi, cần xác lập rõ vai trò của các bên liên quan ở ba cấp độ. Ở cấp vĩ mô, Bộ Y tế có thể là cơ quan đầu mối cho thử nghiệm lâm sàng ngẫu nhiên đa trung tâm thông qua các bệnh viện chuyên khoa tâm thần và viện nghiên cứu y học. Bộ Giáo dục và Đào tạo cùng các Sở Giáo dục đóng vai trò chủ trì việc tích hợp chương trình CBT vào môn học và đào tạo giáo viên chủ nhiệm. Sự phối hợp liên bộ này cần được điều phối qua một hội đồng cấp quốc gia với cơ chế họp định kỳ và báo cáo kết quả minh bạch.'),
    ('P', 'Ở cấp trung mô, các trường đại học sư phạm và trường đại học có ngành tâm lý học chịu trách nhiệm đào tạo nhân lực chuyên môn và phối hợp chuyển giao công nghệ. Các viện nghiên cứu, bệnh viện tâm thần đóng vai trò là đơn vị triển khai thử nghiệm và xây dựng cơ sở bằng chứng. Khu vực tư nhân và các tổ chức phi chính phủ có vai trò bổ sung trong phát triển nền tảng số và cung cấp dịch vụ chăm sóc khi hệ thống công lập quá tải.'),
    ('P', 'Ở cấp vi mô, các trường THCS là điểm tiếp xúc trực tiếp với học sinh, là nơi triển khai chương trình can thiệp cụ thể, đồng thời cũng là nguồn dữ liệu thực địa cho các nghiên cứu. Vai trò của hiệu trưởng và đội ngũ giáo viên chủ nhiệm rất quan trọng trong việc tạo môi trường thuận lợi cho triển khai – bao gồm sắp xếp thời khóa biểu, hỗ trợ vận hành và phối hợp với phụ huynh.'),
    ('P', 'Sự phối hợp ba cấp độ này không thể được hình thành một cách tự phát mà cần được thiết kế và điều phối có chủ ý. Mô hình đề xuất là thành lập một Ban Điều phối Quốc gia về Sức khỏe Tâm thần Học đường, với sự tham gia của đại diện ba Bộ (Y tế, Giáo dục, Khoa học – Công nghệ), các chuyên gia học thuật, và đại diện cộng đồng. Ban này có nhiệm vụ ra các khuyến nghị chiến lược, điều phối nguồn lực và giám sát thực hiện.'),

    ('H2', '5.9. Hàm ý cho các chương tiếp theo của luận án'),
    ('P', 'Tổng quan can thiệp ở chương này có ba hàm ý trực tiếp cho các chương tiếp theo của luận án. Thứ nhất, các chương thực trạng cần thu thập dữ liệu không chỉ về tỷ lệ và yếu tố nguy cơ mà còn về hiện trạng tiếp cận dịch vụ tâm lý của HSTHCS — bao gồm các kênh đã có và các rào cản tiếp cận. Thứ hai, nếu luận án có phần thực nghiệm can thiệp, thiết kế nên dựa trên một trong các giao thức CBT đã được kiểm chứng quốc tế (Coping Cat, FRIENDS, Cool Kids, Journey of the Brave) và được bản địa hóa cẩn thận cho ngữ cảnh Việt Nam, ưu tiên triển khai trong môi trường trường học. Thứ ba, các khuyến nghị cuối luận án nên đề cập rõ tới các hướng nghiên cứu ưu tiên đã nêu ở mục 5.7, đặt trong khung phối hợp ba cấp độ ở mục 5.8.'),
    ('P', 'Mở rộng hơn, tổng quan này gợi ý rằng một luận án nghiên cứu can thiệp cho HSTHCS Việt Nam nên cân nhắc thiết kế dạng kết hợp: vừa khảo sát thực trạng (yếu tố nguy cơ, tỷ lệ, biểu hiện) vừa thử nghiệm một can thiệp mới hoặc bản địa hóa một giao thức quốc tế. Thiết kế kết hợp này tận dụng được sự đầu tư về phương tiện đo lường và nguồn dữ liệu, đồng thời cho phép tạo ra cả tri thức về thực trạng và bằng chứng về hiệu quả can thiệp.'),
]

SEC_KET_LUAN = [
    ('H1', '6. Kết luận chương'),
    ('P', 'Chương này đã hệ thống hóa ba nhóm chính của can thiệp CBT cho rối loạn lo âu ở HSTHCS: CBT đã xác nhận hiệu quả qua thử nghiệm lâm sàng, CBT triển khai tại trường, và CBT qua nền tảng số. Khung phân tích ba kênh × ba tầng được sử dụng xuyên suốt cho phép phân loại các nghiên cứu theo cấp độ dịch vụ và kênh cung cấp, làm rõ chỗ nào đã có bằng chứng quốc tế và chỗ nào còn khoảng trống ở Việt Nam. Tại mỗi nhóm can thiệp, các phân tích tổng hợp quốc tế đã cung cấp ước tính hiệu ứng đáng tin cậy; trong khi đó, ở Việt Nam hầu hết các báo cáo còn ở dạng thử nghiệm pilot quy mô nhỏ hoặc nghiên cứu lý luận chưa được củng cố bằng dữ liệu thực nghiệm đối chứng.'),
    ('P', 'Nhóm tác giả nhận diện năm khoảng trống chính trong nghiên cứu can thiệp tại Việt Nam: (i) thiếu thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu HSTHCS; (ii) thiếu chương trình CBT trường học phù hợp với Chương trình giáo dục phổ thông 2018; (iii) thiếu nền tảng CBT số tiếng Việt được kiểm định; (iv) thiếu các nghiên cứu thiết kế và đo lường kết cục dài hạn; (v) thiếu công cụ sàng lọc và đo lường chuẩn hóa tiếng Việt. Trên cơ sở các khoảng trống này, nhóm tác giả đề xuất năm hướng nghiên cứu ưu tiên cho giai đoạn 2026–2030, đặt trong khung phối hợp ba cấp độ giữa vĩ mô, trung mô và vi mô.'),
    ('P', 'Tổng quan này, kết hợp với chương về yếu tố nguy cơ, cung cấp một bức tranh toàn diện đặt nền móng cho các nghiên cứu can thiệp ưu tiên trong giai đoạn 2026–2030. Sự kết hợp hai chương — một về yếu tố nguy cơ, một về khoảng trống can thiệp — phản ánh logic của một chương trình nghiên cứu dài hạn: nhận diện yếu tố nguy cơ trước, sau đó thiết kế can thiệp phù hợp với các yếu tố đó. Các phát hiện trong chương sẽ được sử dụng để định hình thiết kế nghiên cứu thực nghiệm ở các chương sau của luận án, cũng như khung khuyến nghị cuối cùng cho hệ thống can thiệp sức khỏe tâm thần học đường tại Việt Nam.'),
]

SEC_TLTK = [
    ('H1', '7. Tài liệu tham khảo'),
] + [('P', t) for t in TLTK]


# ============================================================
# BUILD FULL + MEDIUM
# ============================================================
def build_FULL():
    doc = doc_init()
    # Title
    H(doc, 'Khoảng trống can thiệp rối loạn lo âu ở học sinh trung học cơ sở', 0)
    P(doc, 'Tổng quan tài liệu phục vụ luận án tiến sĩ Tâm lý học (bản mở rộng FULL — 26/05/2026)', indent=False)
    doc.add_paragraph()
    # Sections
    for sec in [SEC_MO_DAU, SEC_KHUNG, SEC_PHUONG_PHAP, SEC_KET_QUA, SEC_BAN_LUAN, SEC_KET_LUAN, SEC_TLTK]:
        render(doc, sec)
    out = os.path.join(OUTDIR, 'Bai2_LuanAn_FULL_v1_26052026.docx')
    doc.save(out)
    return out

# MEDIUM: bo SEC_KHUNG, rut gon SEC_PHUONG_PHAP (giu 3.1-3.4), rut gon SEC_BAN_LUAN (bo 5.4, 5.5, 5.9), bo SEC_KET_QUA mục 4.4
SEC_PHUONG_PHAP_MEDIUM = [b for b in SEC_PHUONG_PHAP if not (
    (b[0] == 'H2' and (b[1].startswith('3.5') or b[1].startswith('3.6'))) or
    (b[0] == 'P' and ('Do thiết kế tổng quan tự sự, tổng hợp định lượng' in b[1] or
                      'Tổng hợp định tính được thực hiện theo phương pháp' in b[1] or
                      'Thiết kế tổng quan tự sự có ba hạn chế' in b[1]))
)]

SEC_KET_QUA_MEDIUM = [b for b in SEC_KET_QUA if not (
    (b[0] == 'H2' and b[1].startswith('4.4')) or
    (b[0] == 'P' and ('Ngoài việc triển khai từng kênh riêng lẻ' in b[1] or
                      'Mô hình can thiệp pha trộn (blended intervention)' in b[1] or
                      'Đối với Việt Nam, mô hình stepped care kết hợp blended' in b[1] or
                      'Một thách thức quan trọng đối với mô hình stepped care' in b[1])) or
    (b[0] == 'CAP' and 'Bảng 2' in b[1]) or
    (b[0] == 'TABLE' and 'Hiệu quả (cỡ effect tổng hợp)' in (b[1][0] if b[1] else ''))
)]

SEC_BAN_LUAN_MEDIUM = [b for b in SEC_BAN_LUAN if not (
    (b[0] == 'H2' and (b[1].startswith('5.4') or b[1].startswith('5.5') or b[1].startswith('5.9'))) or
    (b[0] == 'P' and ('Đa số các nghiên cứu can thiệp tại Việt Nam đang sử dụng' in b[1] or
                      'Thiết kế thử nghiệm lâm sàng ngẫu nhiên cụm (cluster RCT)' in b[1] or
                      'Bốn vấn đề về đo lường cần được giải quyết' in b[1] or
                      'Đặc biệt, các nghiên cứu can thiệp số cần báo cáo' in b[1] or
                      'Tổng quan can thiệp ở chương này có ba hàm ý trực tiếp' in b[1] or
                      'Mở rộng hơn, tổng quan này gợi ý rằng một luận án' in b[1]))
)]

def build_MEDIUM():
    doc = doc_init()
    H(doc, 'Khoảng trống can thiệp rối loạn lo âu ở học sinh trung học cơ sở', 0)
    P(doc, 'Tổng quan tài liệu phục vụ luận án tiến sĩ Tâm lý học (bản MEDIUM — 26/05/2026)', indent=False)
    doc.add_paragraph()
    for sec in [SEC_MO_DAU, SEC_PHUONG_PHAP_MEDIUM, SEC_KET_QUA_MEDIUM, SEC_BAN_LUAN_MEDIUM, SEC_KET_LUAN, SEC_TLTK]:
        render(doc, sec)
    out = os.path.join(OUTDIR, 'Bai2_LuanAn_MEDIUM_v1_26052026.docx')
    doc.save(out)
    return out


if __name__ == '__main__':
    f1 = build_FULL()
    f2 = build_MEDIUM()
    # Verify word counts
    from docx import Document as D
    for path in [f1, f2]:
        d = D(path)
        w = sum(len(p.text.split()) for p in d.paragraphs)
        t_count = len(d.tables)
        p_count = sum(1 for p in d.paragraphs if p.text.strip())
        size = os.path.getsize(path) // 1024
        print(f"{os.path.basename(path)}: {w} tu | {p_count} doan | {t_count} bang | {size}KB")





