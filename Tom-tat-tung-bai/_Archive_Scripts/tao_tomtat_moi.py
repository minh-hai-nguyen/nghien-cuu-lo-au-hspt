# -*- coding: utf-8 -*-
"""Tóm tắt CTH v5 cho các bài mới (VN 14-18 + QT 21-29) — ≤2,5 trang/bài"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

def make_doc():
    doc = Document()
    s = doc.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(4); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections: sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5); sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc

def rb(doc, t):
    p = doc.add_paragraph(); r = p.add_run(t); r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED
def bl(doc, t, bold=False):
    p = doc.add_paragraph(); r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold
def rh2(doc, t):
    h = doc.add_heading(t, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED
def bh3(doc, t):
    h = doc.add_heading(t, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE
def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear'); cell._tc.get_or_add_tcPr().append(s)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def rate(doc, stars, text):
    p = doc.add_paragraph()
    r = p.add_run(f'Đánh giá: {stars} {text}')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

papers = [
    # (filename, title, authors_short, journal, n, tool, key_results, gender, critique, gap, rating)
    ('VN14_HoangTrungHoc_2025.docx',
     'Hoàng Trung Học & Nguyễn Thùy Dung (2025)',
     'Mức độ căng thẳng, lo âu và trầm cảm ở VTN trong và sau COVID-19 tại Việt Nam',
     'Am J Psychiatric Rehabilitation, 28(1). 8.473 VTN, 6 tỉnh, DASS-21.',
     [['Căng thẳng', '65,5%', '55,4%', '\u221910,1'],
      ['Lo âu', '41,5%', '25,4%', '\u221916,1'],
      ['Trầm cảm', '34,2%', '20,1%', '\u221914,1'],
      ['Điểm DASS-21 TB', '26,68', '22,07', 'p<0,01']],
     'Quan hệ cha mẹ Beta=0,272 (mạnh nhất). Dùng điện tử Beta=0,176. Giấc ngủ Beta=\u20130,149.',
     'Tạp chí không IF rõ. 2 mẫu khác nhau (không dọc thực). R\u00b2=0,190 chỉ giải thích 19%.',
     'NC dọc thực sự theo cùng nhóm. So DASS-21 vs DISC-5. Phân tích giới tính chi tiết.',
     '\u2b50\u2b50\u2b50 Cỡ mẫu lớn nhất VN (8.473), so sánh trong vs sau COVID, nhưng tạp chí yếu.'),

    ('VN15_NgoAnhVinh_2024.docx',
     'Ngô Anh Vinh và cộng sự (2024)',
     'Sức khỏe tâm thần ở VTN dân tộc thiểu số tại Việt Nam — Lạng Sơn',
     'J Affective Disorders Reports, 17:100795. 845 HS DTTS nội trú, DASS-21 + ACEs.',
     [['Trầm cảm', '59,0%', '', ''],
      ['Lo âu', '54,4%', '', ''],
      ['Căng thẳng', '24,7%', '', ''],
      ['ACEs (\u22651)', '48,9%', 'TB=1,1', 'SD=1,8']],
     'ACEs \u2192 lo âu Coef.=0,28. Bạn bè kém \u2192 trầm cảm OR=6,84. Sống \u226510 người bảo vệ OR=0,51.',
     'Chỉ 1 tỉnh (Lạng Sơn). DASS-21 ngưỡng thấp. Không phân tích giới tính chi tiết.',
     'Mở rộng nhiều tỉnh vùng cao. So sánh DTTS nội trú vs không nội trú. Can thiệp bạn bè.',
     '\u2b50\u2b50\u2b50 NC đầu tiên DTTS nội trú VN, DASS-21+ACEs, Elsevier OA.'),

    ('VN16_BaoQuyen_2025.docx',
     'Nguyễn Ngọc Bảo Quyên và cộng sự (2025)',
     'Thực trạng SKTT HS THPT Hà Nội — trầm cảm, lo âu, căng thẳng',
     'TC Y học Cộng đồng, 66, CD10, tr. 79\u201386. 501 HS, 8\u201312/2023, DASS-21.',
     [['Trầm cảm', '78,8%', 'Nặng+RN', '33,2%'],
      ['Lo âu', '86,2%', 'Nặng+RN', '54,5%'],
      ['Căng thẳng', '76,6%', 'Nặng+RN', '38,3%']],
     'Nữ > nam p<0,05 cả 3 chỉ số. Khối lớp và loại trường không ý nghĩa.',
     'Lo âu 86,2% CỰC CAO — ngưỡng DASS thấp. Nữ 78% mẫu — thiên lệch. Online survey. n=501 nhỏ.',
     'Cỡ mẫu lớn hơn, tỷ lệ giới cân bằng. So DASS vs GAD-7 cùng mẫu.',
     '\u2b50\u2b50 Đề tài thời sự nhưng tỷ lệ cao bất thường, thiên lệch giới và tự chọn.'),

    ('VN17_DanhLam_2022.docx',
     'Nguyễn Danh Lâm và cộng sự (2022)',
     'Stress, lo âu, trầm cảm HS THPT Yên Định, Thanh Hóa',
     'TC Y học Việt Nam, 516(1), tr. 67\u201370. 482 HS, 9/2021, DASS-21.',
     [['Stress', '41,7%', '', ''],
      ['Lo âu', '49,0%', '', ''],
      ['Trầm cảm', '43,6%', '', ''],
      ['Tự làm đau', '10,0%', 'Nghĩ tự tử', '23,6%']],
     'Tự hại 10%, cố tự tử 1,4% — đáng báo động. Chủ yếu mức nhẹ-vừa.',
     'Chọn mẫu có chủ đích (2 trường). Chỉ mô tả, không hồi quy. Bài 4 trang.',
     'Phân tích yếu tố nguy cơ (hồi quy). Can thiệp giảm tự hại.',
     '\u2b50\u2b50 Vùng bán đô thị, dữ liệu tự hại quan trọng nhưng bài quá ngắn.'),

    ('VN18_AnGiang_2025.docx',
     'Lê Minh T., Nguyễn Đăng K., Ngô Anh V. (2025)',
     'Sàng lọc lo âu, trầm cảm, stress HS THPT Long Bình, An Giang',
     'TC Y học Việt Nam, 549(1), tr. 32\u201335. 366 HS, 6/2024, DASS-21.',
     [['Lo âu', '61,2%', '', ''],
      ['Trầm cảm', '47,3%', '', ''],
      ['Stress', '38,0%', '', ''],
      ['Đồng mắc cả 3', '27,3%', '', '']],
     'Nam/nữ cân bằng (50,5/49,5%). 27,3% đồng mắc cả 3 triệu chứng.',
     '1 trường duy nhất. Chỉ mô tả. Bài ngắn.',
     'Mở rộng ĐBSCL. Phân tích hồi quy. So sánh ĐBSCL vs đô thị.',
     '\u2b50\u2b50 Vùng ĐBSCL ít NC, giới cân bằng, nhưng 1 trường và chỉ mô tả.'),
]

qt_papers = [
    ('QT21_Norway_2025.docx',
     'Norway mental distress trends 2011\u20132024',
     'Xu hướng tăng căng thẳng tâm thần ở VTN Na Uy 13 năm',
     'Social Science & Medicine, 2025 (Q1, IF\u22485,4). 10 trang.',
     'Bất mãn trường học và mạng xã hội giải thích xu hướng tăng.',
     '\u2b50\u2b50\u2b50\u2b50 Q1, xu hướng 13 năm, phân tích decomposition.'),

    ('QT22_ScreenTime_2025.docx',
     'Screen time & depression/anxiety longitudinal',
     'Mối liên quan dọc giữa thời gian màn hình và trầm cảm/lo âu VTN',
     'British J Clinical Psychology, 2025, 64:873\u2013887 (Q1). 15 trang.',
     'Nghiên cứu DỌC — screen time liên quan trầm cảm/lo âu.',
     '\u2b50\u2b50\u2b50\u2b50 Thiết kế dọc vượt trội. BJCP Q1.'),

    ('QT23_JAACAP_US_2024.docx',
     'US trends mental disorders 2013\u20132021 (JAACAP)',
     'Xu hướng rối loạn tâm thần trẻ em/VTN trong hệ thống y tế công Mỹ',
     'JAACAP, 2024 (Q1, IF\u224811,0). 15 trang.',
     'Xu hướng 8 năm trong hệ thống quốc gia Mỹ.',
     '\u2b50\u2b50\u2b50\u2b50\u2b50 JAACAP Q1 IF=11, dữ liệu quốc gia Mỹ.'),

    ('QT24_WHO_Europe_2025.docx',
     'WHO Europe — SKTT trẻ em/thanh niên (Lancet)',
     'Tổng quan SKTT trẻ em/thanh niên khu vực WHO châu Âu',
     'Lancet Regional Health Europe, 2025 (Q1, IF\u224815,0). 13 trang.',
     '9 triệu VTN châu Âu có rối loạn SKTT. Lo âu+trầm cảm >50% ca.',
     '\u2b50\u2b50\u2b50\u2b50\u2b50 Lancet Q1 IF=15, phạm vi toàn châu Âu.'),

    ('QT26_UK_NHS_2025.docx',
     'UK NHS Mental Health Statistics 2025',
     'Thống kê SKTT Anh: tỷ lệ, dịch vụ, tài trợ',
     'House of Commons Library, UK Parliament, 2025. 46 trang.',
     '25,8% thanh niên 16\u201324 rối loạn TT (tăng từ 18,9% 2014). Nữ 36,1% vs nam 16,3%.',
     '\u2b50\u2b50\u2b50 Nguồn chính thức Quốc hội Anh, nhưng báo cáo chính sách không NC gốc.'),

    ('QT27_Nature_SocialMedia_2025.docx',
     'Social media use in adolescents (Nature Human Behaviour)',
     'Mạng xã hội ở VTN có/không rối loạn SKTT — UK n=3.340',
     'Nature Human Behaviour, 9, 2025, pp. 1283\u20131299 (Q1, IF\u224824,0). 21 trang.',
     'VTN có rối loạn SKTT dùng MXH nhiều hơn, ít hài lòng hơn. Chẩn đoán lâm sàng.',
     '\u2b50\u2b50\u2b50\u2b50\u2b50 Nature Q1 IF=24, mẫu quốc gia UK, chẩn đoán lâm sàng.'),

    ('QT28_AJP_Treatment_2024.docx',
     'Pediatric anxiety treatment (AJP)',
     'Phương pháp hiện tại và tương lai điều trị rối loạn lo âu trẻ em',
     'American J Psychiatry, 2024 (Q1, IF\u224818,0). 12 trang.',
     'CBT 47\u201366% phục hồi. SSRI. Hướng mới: digital, neuromodulation.',
     '\u2b50\u2b50\u2b50\u2b50\u2b50 AJP Q1 IF=18, tổng quan toàn diện cập nhật nhất.'),

    ('QT29_CBT_NetworkMeta_2025.docx',
     'CBT interventions network meta-analysis (BMC Psychiatry)',
     'So sánh hiệu quả can thiệp lo âu trẻ em — 30 RCTs, 1.711 HS',
     'BMC Psychiatry, 2025 (Q1, OA). 14 trang.',
     'Network meta-analysis Bayesian so sánh CBT, thuốc, kết hợp, tâm lý giáo dục.',
     '\u2b50\u2b50\u2b50\u2b50 BMC Q1, 30 RCTs, phương pháp nâng cao nhất.'),
]

# Generate VN summaries
for fname, authors, title, journal, results, findings, critique, gap, rating in papers:
    d = make_doc()
    bl(d, f'Tóm tắt: {authors}', bold=True)
    rb(d, 'Tên công trình, tác giả, năm, mẫu khảo sát')
    bl(d, f'{title}. {journal}')

    rb(d, 'Kết quả nghiên cứu định lượng')
    if len(results[0]) == 4:
        h = results[0]
        if results[0][2]:  # has 3rd column
            tbl(d, ['Tình trạng', 'Trong COVID / Tổng', 'Sau COVID / Chi tiết', 'Thay đổi / %'],
                results, widths=[3.0, 3.0, 3.0, 2.5])
        else:
            tbl(d, ['Tình trạng', 'Tỷ lệ', '', ''],
                results, widths=[3.5, 2.5, 2.5, 2.5])
    d.add_paragraph()

    rb(d, 'Nhận xét')
    bl(d, findings)

    rb(d, 'Kết luận')
    bl(d, f'{title}. {findings[:150]}...', bold=True)

    rh2(d, 'Phản biện')
    bl(d, critique)
    rh2(d, 'Hướng nghiên cứu tiếp theo')
    bl(d, gap)

    p = d.add_paragraph()
    r = p.add_run(f'Đánh giá: {rating}')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    d.save(fname)
    sys.stderr.write(f'{fname} OK\n')

# Generate QT summaries
for fname, title_short, title_vn, journal, findings, rating in qt_papers:
    d = make_doc()
    bl(d, f'Tóm tắt: {title_short}', bold=True)
    rb(d, 'Tên công trình')
    bl(d, f'{title_vn}. {journal}')

    rb(d, 'Phát hiện chính')
    bl(d, findings)

    rh2(d, 'Phản biện')
    bl(d, f'Xem bản dịch chi tiết trong thư mục 03_Ban-dich/.')

    p = d.add_paragraph()
    r = p.add_run(f'Đánh giá: {rating}')
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

    d.save(fname)
    sys.stderr.write(f'{fname} OK\n')
