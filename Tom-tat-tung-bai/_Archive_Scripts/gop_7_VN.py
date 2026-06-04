# -*- coding: utf-8 -*-
"""Gộp 7 tóm tắt VN14-20 vào 1 file"""
import os, sys
os.environ['PYTHONIOENCODING'] = 'utf-8'
import docx
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_BREAK

OUT = os.path.dirname(os.path.abspath(__file__))

files = [
    ('VN14_HoangTrungHoc_2025.docx', 'VN14 — Hoàng Trung Học & Nguyễn Thùy Dung (2025)'),
    ('VN15_NgoAnhVinh_2024.docx', 'VN15 — Ngô Anh Vinh và cs. (2024)'),
    ('VN16_BaoQuyen_2025.docx', 'VN16 — Nguyễn Ngọc Bảo Quyên và cs. (2025)'),
    ('VN17_DanhLam_2022.docx', 'VN17 — Nguyễn Danh Lâm và cs. (2022)'),
    ('VN18_AnGiang_2025.docx', 'VN18 — Lê Minh T., Nguyễn Đăng K., Ngô Anh V. (2025) — Long Bình, An Giang'),
    ('VN19_TranThaoVi_2025.docx', 'VN19 — Hồ Thị Trúc Quỳnh (2025) — Huế (lưu ý: "Trần Thảo Vi" trong đề cương gốc là nhầm lẫn)'),
    ('VN20_TranHoVinhLoc_2024.docx', 'VN20 — Trần Hồ Vĩnh Lộc và cs. (2024) — TPHCM'),
]

# Create merged doc with same style
merged = docx.Document()
style = merged.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.line_spacing = 1.5
for s in merged.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2)

# Cover page
h = merged.add_heading('TÓM TẮT 7 BÀI NGHIÊN CỨU VIỆT NAM (VN14–VN20)', level=1)
for r in h.runs:
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0, 0, 0)

p = merged.add_paragraph()
r = p.add_run('Lo âu, trầm cảm, căng thẳng ở học sinh THCS/THPT — Việt Nam, 2022–2025')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(13)

p = merged.add_paragraph()
r = p.add_run('Gộp ngày 05/04/2026. 7 bài nghiên cứu theo chuẩn Công Thị Hằng v5.')
r.font.name = 'Times New Roman'; r.italic = True

# Table of contents
merged.add_paragraph()
h = merged.add_heading('MỤC LỤC', level=2)
for r in h.runs:
    r.font.name = 'Times New Roman'
    r.font.color.rgb = RGBColor(0, 0, 0)

for filename, label in files:
    p = merged.add_paragraph()
    r = p.add_run(f'• {label}')
    r.font.name = 'Times New Roman'

merged.add_page_break()

# Helper to copy element
from copy import deepcopy

# Merge each file
for filename, label in files:
    filepath = os.path.join(OUT, filename)
    if not os.path.exists(filepath):
        print(f'MISSING: {filename}')
        continue

    src = docx.Document(filepath)

    # Add section header
    h = merged.add_heading(label, level=1)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

    # Copy all body elements (paragraphs + tables) in order
    for element in src.element.body:
        if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
            # Clone the element
            new_elem = deepcopy(element)
            merged.element.body.append(new_elem)

    merged.add_page_break()
    print(f'OK: {filename}')

# Remove last page break
if merged.paragraphs and 'sectPr' not in merged.paragraphs[-1]._p.xml:
    pass

outpath = os.path.join(OUT, '..', 'Tom tat 7 bai VN - 05042026.docx')
merged.save(outpath)

# Verify
d = docx.Document(outpath)
t = '\n'.join([p.text for p in d.paragraphs])
print(f'\nSaved: {outpath}')
print(f'  Total: {len(t)} chars, {len(d.tables)} tables, {len([p for p in d.paragraphs if p.text.strip()])} paras')
