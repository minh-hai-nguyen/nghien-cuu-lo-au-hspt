"""Bulk fix 47 NO_REFS summaries by appending a 'Tham khao' section
auto-filled from `02_Papers-goc/canonical_index.json` (filename-level metadata only).

Per `feedback_verify_tung_entry_truoc_khi_gui.md`: we DO NOT fabricate journal
names, full author lists, volume/issue/pages, DOI/PMID. Fields not present in
canonical_index are explicitly marked "[khong co trong canonical]" so the
downstream reviewer knows they still need manual verification against the PDF
goc + Crossref/PubMed.

Output: modify file IN PLACE (per task spec).
Strip metadata: also strip core_properties (title, author, comments, etc.).
"""
from __future__ import annotations

import json
import re
import sys
from datetime import date
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

ROOT = Path(r"c:/Users/OS/OneDrive/read_books/Lo-au")
TOM_TAT_DIR = ROOT / "Tom-tat-tung-bai"
CANONICAL = ROOT / "02_Papers-goc" / "canonical_index.json"
RESULTS = ROOT / "06_Scripts" / "_audit_missing_refs_results.json"
REPORT_MD = ROOT / "06_Scripts" / "BULK_FIX_NoRefs_RESULT_07062026.md"

TODAY = date(2026, 6, 7).isoformat()
FOOTER = f"Bo sung Tham khao 07/06/2026 (auto-fill tu canonical_index)"
HEADING = "Tham khao"
MISSING = "[khong co trong canonical]"

SKIP_FILE_PATTERNS = [
    "_FIXED_07062026",
    "_FIXED_27052026",
]


def parse_id_from_filename(name: str) -> str | None:
    m = re.match(r"^(QT\d{3}|VN\d{3})", name)
    return m.group(1) if m else None


def parse_year(descriptor: str) -> str | None:
    m = re.search(r"(19|20)\d{2}", descriptor or "")
    return m.group(0) if m else None


def parse_lead_author(descriptor: str) -> str:
    """Heuristic: descriptor like 'Jenkins_et_al_2023_USA' -> 'Jenkins et al.'
    'Mandaknalli_Malusare_2021' -> 'Mandaknalli & Malusare'
    'NSCH_2020_USA' -> 'NSCH' (org)
    'VNAMHS_2022_National' -> 'VNAMHS'
    """
    if not descriptor:
        return MISSING
    parts = descriptor.split("_")
    # take tokens until we hit a year
    name_tokens: list[str] = []
    for t in parts:
        if re.match(r"^(19|20)\d{2}$", t):
            break
        name_tokens.append(t)
    if not name_tokens:
        return descriptor
    # 'et_al' handling
    if "et" in name_tokens and "al" in name_tokens:
        before = []
        for t in name_tokens:
            if t in {"et", "al"}:
                break
            before.append(t)
        return (" ".join(before) + " et al.").strip()
    # two-author pattern (no et_al)
    if len(name_tokens) == 2:
        return f"{name_tokens[0]} & {name_tokens[1]}"
    # one token
    if len(name_tokens) == 1:
        return name_tokens[0]
    # more than two -> join as-is
    return " ".join(name_tokens)


def build_pdf_path(entry: dict) -> str:
    pdf = entry.get("pdf_path")
    if pdf:
        return pdf
    folder = entry.get("pdf_folder")
    pdf_name = entry.get("pdf")
    if folder and pdf_name:
        return f"02_Papers-goc/{folder}/{pdf_name}"
    if pdf_name:
        return f"02_Papers-goc/{pdf_name}"
    return MISSING


def build_apa(entry: dict) -> str:
    descriptor = entry.get("descriptor") or ""
    author = parse_lead_author(descriptor)
    year = parse_year(descriptor) or MISSING
    return (
        f"{author} ({year}). [Tieu de bai bao: {MISSING}]. "
        f"[Tap chi/Nguon: {MISSING}], [vol/issue/pages: {MISSING}]."
    )


def style_para(p, *, bold: bool = False, size: int = 12):
    for r in p.runs:
        r.font.name = "Times New Roman"
        r.font.size = Pt(size)
        if bold:
            r.bold = True
    pf = p.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_para(doc, text: str, *, bold: bool = False, size: int = 12):
    p = doc.add_paragraph(text)
    style_para(p, bold=bold, size=size)
    return p


def append_refs_section(docx_path: Path, entry: dict) -> None:
    doc = Document(str(docx_path))

    # spacer
    doc.add_paragraph("")

    # Heading
    add_para(doc, HEADING, bold=True, size=13)

    # APA-7 citation
    add_para(doc, build_apa(entry))

    # PDF path
    add_para(doc, f"PDF path: {build_pdf_path(entry)}")

    # DOI / PMID
    add_para(doc, f"DOI: {MISSING} | PMID: {MISSING}")

    # blank + footer
    doc.add_paragraph("")
    add_para(doc, FOOTER)

    # Strip metadata (core properties)
    cp = doc.core_properties
    cp.author = ""
    cp.last_modified_by = ""
    cp.title = ""
    cp.subject = ""
    cp.keywords = ""
    cp.comments = ""
    cp.category = ""
    cp.content_status = ""

    doc.save(str(docx_path))


def main():
    sys.stdout.reconfigure(encoding="utf-8")

    canonical = json.loads(CANONICAL.read_text(encoding="utf-8"))
    audit = json.loads(RESULTS.read_text(encoding="utf-8"))

    # Filter NO_REFS only
    no_refs = [r for r in audit if r["classification"] == "NO_REFS"]

    fixed: list[dict] = []
    skipped: list[dict] = []
    sample_before = None
    sample_after = None

    for rec in no_refs:
        fname = rec["file"]

        # Skip rules
        if any(pat in fname for pat in SKIP_FILE_PATTERNS):
            skipped.append({"file": fname, "reason": "FIXED_DATE_TAG"})
            continue

        path = TOM_TAT_DIR / fname
        if not path.exists():
            skipped.append({"file": fname, "reason": "FILE_NOT_FOUND"})
            continue

        # archive/bai-bao-khgdvn already excluded by audit scope, but guard:
        if "_Archive" in str(path) or "bai-bao-khgdvn" in str(path):
            skipped.append({"file": fname, "reason": "EXCLUDED_DIR"})
            continue

        cid = parse_id_from_filename(fname)
        if not cid or cid not in canonical:
            skipped.append({"file": fname, "reason": "ORPHAN_NO_CANONICAL"})
            continue

        entry = canonical[cid]

        # Capture sample BEFORE on the first file
        if sample_before is None:
            try:
                d0 = Document(str(path))
                tail = [p.text for p in d0.paragraphs[-6:]]
                sample_before = {"file": fname, "tail": tail}
            except Exception:
                pass

        try:
            append_refs_section(path, entry)
            fixed.append({
                "file": fname,
                "id": cid,
                "pdf_path": build_pdf_path(entry),
                "apa": build_apa(entry),
            })

            if sample_after is None and sample_before and sample_before["file"] == fname:
                d1 = Document(str(path))
                tail = [p.text for p in d1.paragraphs[-10:]]
                sample_after = {"file": fname, "tail": tail}
        except Exception as exc:  # pragma: no cover - defensive
            skipped.append({"file": fname, "reason": f"WRITE_ERROR: {exc}"})

    # Write report
    lines: list[str] = []
    lines.append("# BULK FIX — 47 NO_REFS summaries (07/06/2026)")
    lines.append("")
    lines.append(f"**Script:** `06_Scripts/_bulk_fix_norefs_07062026.py`")
    lines.append(f"**Source audit:** `06_Scripts/_audit_missing_refs_results.json`")
    lines.append(f"**Canonical source:** `02_Papers-goc/canonical_index.json`")
    lines.append(f"**Run date:** {TODAY}")
    lines.append("")
    lines.append("## 1. Totals")
    lines.append("")
    lines.append(f"- Total NO_REFS processed: **{len(no_refs)}**")
    lines.append(f"- Successfully fixed (in place): **{len(fixed)}**")
    lines.append(f"- Skipped: **{len(skipped)}**")
    lines.append("")
    lines.append("## 2. Skipped detail")
    lines.append("")
    if skipped:
        for s in skipped:
            lines.append(f"- `{s['file']}` — {s['reason']}")
    else:
        lines.append("_None._")
    lines.append("")
    lines.append("## 3. Fixed list")
    lines.append("")
    for f in fixed:
        lines.append(f"- `{f['file']}` (ID={f['id']}) -> PDF: `{f['pdf_path']}`")
    lines.append("")
    lines.append("## 4. Sample diff (1 example)")
    lines.append("")
    if sample_before:
        lines.append(f"**File:** `{sample_before['file']}`")
        lines.append("")
        lines.append("**Last 6 paragraphs BEFORE:**")
        lines.append("")
        lines.append("```")
        for t in sample_before["tail"]:
            lines.append(t or "<empty>")
        lines.append("```")
        lines.append("")
    if sample_after:
        lines.append("**Last 10 paragraphs AFTER:**")
        lines.append("")
        lines.append("```")
        for t in sample_after["tail"]:
            lines.append(t or "<empty>")
        lines.append("```")
        lines.append("")
    lines.append("## 5. Notes / IMPORTANT caveats")
    lines.append("")
    lines.append(
        "- canonical_index.json chỉ chứa **metadata tệp tin** (descriptor, pdf, folder); "
        "KHÔNG có authors đầy đủ / journal / volume / issue / pages / DOI / PMID."
    )
    lines.append(
        "- Phần Tham khảo auto-fill này chỉ thiết lập **khung tối thiểu** (lead author, year, PDF path) "
        f"và đánh dấu các trường còn lại bằng `{MISSING}`."
    )
    lines.append(
        "- Vẫn cần **bước tiếp theo (manual)**: mở từng PDF gốc, verify DOI/PMID qua Crossref + PubMed, "
        "điền tiêu đề / tạp chí / vol-issue-pages thật, theo `feedback_verify_tung_entry_truoc_khi_gui.md`."
    )
    lines.append(
        "- Các file đã có `_FIXED_07062026` hoặc `_FIXED_27052026` được giữ nguyên (không bị ghi đè)."
    )

    REPORT_MD.write_text("\n".join(lines), encoding="utf-8")

    print(f"FIXED: {len(fixed)}")
    print(f"SKIPPED: {len(skipped)}")
    print(f"Report -> {REPORT_MD}")


if __name__ == "__main__":
    main()
