"""Doc trả lời: F-statistic là gì + tại sao có F nhỏ - format mới (câu trả lời xanh trước phụ lục)."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/F_statistic_ANOVA_dien_giai.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
BLUE = RGBColor(0, 112, 192)
DARK = RGBColor(31, 73, 125)
GREEN = RGBColor(54, 95, 44)
RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(191, 97, 14)
GRAY = RGBColor(90, 90, 90)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def vn_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p
def blue_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = BLUE
    return p

# ===== TITLE =====
title = d.add_heading('F-statistic (kiểm định F trong ANOVA): F là gì + tại sao có lúc F lại nhỏ?', level=0)
for r in title.runs: r.font.color.rgb = DARK
sub = d.add_paragraph()
sr = sub.add_run('Diễn giải F-statistic qua ví dụ bảng so sánh nam vs nữ về 4 dạng RLLA')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = GRAY
d.add_paragraph()

# ===== CÂU HỎI =====
qbox = d.add_table(rows=1, cols=1); qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]; shade(cell, 'FFF8DC')
p = cell.paragraphs[0]
r = p.add_run('Câu hỏi của thầy:'); r.bold = True
p2 = cell.add_paragraph()
p2.add_run('"Em giải thích giúp thầy F là gì trong ảnh nhé. Tại sao trong bảng có lúc F lại có giá trị nhỏ vậy?"')
p3 = cell.add_paragraph()
r3 = p3.add_run('Bảng có 4 giá trị F: F = 44,484 (p<0,001) | F = 0,246 (p=0,620) | F = 45,984 (p<0,001) | F = 29,642 (p<0,001)')
r3.italic = True; r3.font.size = Pt(10); r3.font.color.rgb = GRAY
d.add_paragraph()

# ===== BỐI CẢNH =====
add_h('Bối cảnh', 1)
vn_para('Bảng so sánh điểm trung bình (ĐTB) của 4 dạng rối loạn lo âu (RLLA) giữa nam (n=614) và nữ (n=738) HS:')
ctx_tbl = d.add_table(rows=3, cols=5); ctx_tbl.style = 'Table Grid'
ctx_hdr = ['', 'RLLATQ', 'RLLACL', 'RLLAXH', 'RLLA tổng']
for i, h in enumerate(ctx_hdr):
    c = ctx_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
ctx_data = [
    ('Nam (n=614)', '51,43 (22,01)', '25,42 (25,46)', '43,20 (25,09)', '40,02 (19,02)'),
    ('Nữ (n=738)', '59,47 (22,07)', '24,76 (23,29)', '52,74 (26,31)', '45,66 (18,91)'),
]
for i, row in enumerate(ctx_data):
    for j, v in enumerate(row):
        ctx_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in ctx_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True

vn_para('(ĐTB = Điểm trung bình; số trong ngoặc = ĐLC = Độ lệch chuẩn)', size=10)
d.add_paragraph()

# ===== KIẾN THỨC NỀN =====
add_h('Kiến thức nền', 1)

vn_para('1. F-statistic là gì?', bold=True, size=12)
vn_para('F = giá trị thống kê của kiểm định ANOVA (Analysis of Variance — Phân tích phương sai). '
        'Dùng để so sánh trung bình của ≥ 2 nhóm.')
vn_para('Công thức: F = MS_between / MS_within', bold=True, size=12)
vn_para('Trong đó:')
for it in [
    'MS_between (Mean Square between groups) = phương sai GIỮA các nhóm — đo CHÊNH LỆCH trung bình giữa nhóm',
    'MS_within (Mean Square within groups) = phương sai TRONG mỗi nhóm — đo BIẾN THIÊN giữa các cá nhân trong cùng 1 nhóm',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)

d.add_paragraph()

vn_para('2. Logic đọc F:', bold=True, size=12)
logic_tbl = d.add_table(rows=4, cols=2); logic_tbl.style = 'Table Grid'
logic_hdr = ['Tình huống', 'F']
for i, h in enumerate(logic_hdr):
    c = logic_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
logic_data = [
    ('Các nhóm có trung bình GẦN BẰNG NHAU (chênh nhỏ)', 'F NHỎ (gần 0) → KHÔNG khác biệt'),
    ('Các nhóm có trung bình KHÁC NHAU rõ rệt (chênh lớn)', 'F LỚN → CÓ khác biệt'),
    ('Phương sai TRONG nhóm rất lớn (data scatter rộng)', 'F nhỏ đi (vì mẫu số tăng)'),
]
for i, row in enumerate(logic_data):
    for j, v in enumerate(row):
        logic_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in logic_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('3. Ngưỡng F (chung):', bold=True, size=12)
th_tbl = d.add_table(rows=5, cols=2); th_tbl.style = 'Table Grid'
th_hdr = ['Khoảng F', 'Ý nghĩa']
for i, h in enumerate(th_hdr):
    c = th_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
th_data = [
    ('F < 1', 'Phương sai trong nhóm > giữa nhóm → KHÔNG có khác biệt rõ'),
    ('F = 1-3', 'Khác biệt nhỏ — thường p > 0,05 (không có ý nghĩa)'),
    ('F = 3-10', 'Khác biệt vừa — có thể có ý nghĩa thống kê (p < 0,05)'),
    ('F > 10', 'Khác biệt rõ ràng — thường p < 0,001'),
]
for i, row in enumerate(th_data):
    for j, v in enumerate(row):
        th_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in th_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('Lưu ý: ngưỡng cụ thể "có ý nghĩa" phụ thuộc vào DEGREES OF FREEDOM (df) — '
        'với df lớn (mẫu lớn) chỉ cần F nhỏ cũng có thể p < 0,05. Đọc F kèm p-value.', bold=False, size=10)
d.add_paragraph()

vn_para('4. Phân tích 4 giá trị F trong bảng', bold=True, size=12)
ana_tbl = d.add_table(rows=5, cols=5); ana_tbl.style = 'Table Grid'
ana_hdr = ['Cột', 'Nam ĐTB', 'Nữ ĐTB', 'Chênh', 'F (p)']
for i, h in enumerate(ana_hdr):
    c = ana_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
ana_data = [
    ('RLLATQ', '51,43', '59,47', '8,04', 'F = 44,484; p < 0,001'),
    ('RLLACL', '25,42', '24,76', '0,66', 'F = 0,246; p = 0,620 ⚠'),
    ('RLLAXH', '43,20', '52,74', '9,54', 'F = 45,984; p < 0,001'),
    ('RLLA tổng', '40,02', '45,66', '5,64', 'F = 29,642; p < 0,001'),
]
for i, row in enumerate(ana_data):
    for j, v in enumerate(row):
        ana_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in ana_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('5. Tính effect size (Cohen\'s d) ngầm trong bảng', bold=True, size=12)
vn_para('Để hỗ trợ diễn giải F, em tính Cohen\'s d ≈ chênh ĐTB / ĐLC trung bình:')
es_tbl = d.add_table(rows=5, cols=3); es_tbl.style = 'Table Grid'
es_hdr = ['Cột', 'Cohen d ước tính', 'Phân loại']
for i, h in enumerate(es_hdr):
    c = es_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255, 255, 255)
es_data = [
    ('RLLATQ', '8,04 / 22 ≈ 0,36', 'Small-medium'),
    ('RLLACL', '0,66 / 24 ≈ 0,03', 'NEGLIGIBLE (gần như không có)'),
    ('RLLAXH', '9,54 / 26 ≈ 0,37', 'Small-medium'),
    ('RLLA tổng', '5,64 / 19 ≈ 0,30', 'Small'),
]
for i, row in enumerate(es_data):
    for j, v in enumerate(row):
        es_tbl.rows[i+1].cells[j].text = v
        if j == 0:
            for pp in es_tbl.rows[i+1].cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('→ Cohen d xác nhận F: RLLACL có effect size 0,03 (NEGLIGIBLE) nên F = 0,246 (rất nhỏ) là HOÀN TOÀN NHẤT QUÁN.')
d.add_paragraph()

# =====================================================
# CÂU TRẢ LỜI (TÔ XANH, GOM LẠI TRƯỚC PHỤ LỤC)
# =====================================================
add_h('Câu trả lời', 1, BLUE)

blue_para('Hỏi 1: F là gì?', bold=True, size=12)
blue_para('F = F-statistic của kiểm định ANOVA (Phân tích phương sai). Công thức: '
          'F = phương sai GIỮA nhóm / phương sai TRONG nhóm. '
          'F dùng để kiểm tra: trung bình các nhóm có khác nhau không?')
d.add_paragraph()

blue_para('Hỏi 2: Tại sao có lúc F lại nhỏ (F = 0,246)?', bold=True, size=12)
blue_para('Vì F PHẢN ÁNH ĐỘ LỚN của khác biệt:', bold=True)
ans_list = [
    '2 nhóm có ĐTB GẦN BẰNG NHAU → tử số (chênh giữa nhóm) ≈ 0 → F ≈ 0',
    '2 nhóm có ĐTB CHÊNH LỆCH LỚN → tử số lớn → F lớn',
    'F = 0,246 KHÔNG phải lỗi tính toán — là KẾT QUẢ THẬT phản ánh dữ liệu',
]
for it in ans_list:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Cụ thể trong bảng:', bold=True, size=12)
ans_tbl = d.add_table(rows=5, cols=3); ans_tbl.style = 'Table Grid'
ans_hdr = ['F', 'Cột', 'Lý do']
for i, h in enumerate(ans_hdr):
    c = ans_tbl.rows[0].cells[i]; shade(c, 'BDD7EE')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = BLUE
ans_data = [
    ('F = 44,484', 'RLLATQ', 'Nam 51,43 vs Nữ 59,47 — chênh 8 điểm → F LỚN'),
    ('F = 0,246', 'RLLACL', 'Nam 25,42 ≈ Nữ 24,76 — chênh CHỈ 0,66 điểm → F NHỎ'),
    ('F = 45,984', 'RLLAXH', 'Nam 43,20 vs Nữ 52,74 — chênh 9,5 điểm → F LỚN'),
    ('F = 29,642', 'RLLA tổng', 'Nam 40,02 vs Nữ 45,66 — chênh 5,6 điểm → F khá LỚN'),
]
for i, row in enumerate(ans_data):
    for j, v in enumerate(row):
        c = ans_tbl.rows[i+1].cells[j]
        pp = c.paragraphs[0]
        rr = pp.add_run(v); rr.font.color.rgb = BLUE; rr.bold = (j == 0)
d.add_paragraph()

blue_para('Diễn giải lâm sàng:', bold=True, size=12)
for it in [
    'RLLATQ, RLLAXH, RLLA tổng (F lớn + p < 0,001): nữ HS lo âu CAO HƠN nam có ý nghĩa thống kê',
    'RLLACL (F = 0,246 + p = 0,620): nam và nữ KHÔNG khác biệt — RLLACL không thiên về giới',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Hàm ý cho can thiệp:', bold=True, size=12)
blue_para('• Chương trình giảm RLLACL: KHÔNG cần phân theo giới, áp dụng cho cả nam lẫn nữ', bold=False)
blue_para('• Chương trình giảm RLLATQ + RLLAXH: ƯU TIÊN nữ HS (vì chênh trên 8 điểm)', bold=False)
d.add_paragraph()

blue_para('Ghi nhớ 1 dòng:', bold=True, size=12)
blue_para('F LỚN = nhóm KHÁC NHAU rõ. F NHỎ (gần 0) = nhóm NHƯ NHAU. '
          'KHÔNG có "F sai" — F luôn phản ánh dữ liệu thực. Đọc F kèm p-value.', bold=True)

d.add_paragraph()

# =====================================================
# PHỤ LỤC
# =====================================================
add_h('Phụ lục — Tài liệu tham khảo', 1)
refs = [
    'Fisher RA. (1925). Statistical Methods for Research Workers. Oliver and Boyd. [Bài gốc giới thiệu F-statistic + ANOVA]',
    'Cohen J. (1988). Statistical Power Analysis for the Behavioral Sciences (2nd ed.). Routledge. [Cohen\'s d: 0,2 small / 0,5 medium / 0,8 large]',
    'Field A. (2018). Discovering Statistics Using IBM SPSS Statistics (5th ed.). Sage. [Giáo khoa ANOVA chuẩn]',
    'Maxwell SE, Delaney HD, Kelley K. (2018). Designing Experiments and Analyzing Data (3rd ed.). Routledge. [Diễn giải F + post-hoc tests]',
    'Doc liên quan trong dự án: 01_Bao-cao/AOR_Adjusted_Odds_Ratio_cach_tinh.docx (regression đa biến); 01_Bao-cao/NNT_va_MannWhitney_chenh_lech_trung_vi.docx (so sánh test khác).',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs: r.font.size = Pt(10)

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
