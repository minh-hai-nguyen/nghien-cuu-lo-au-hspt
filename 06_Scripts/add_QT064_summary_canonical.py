"""Build QT064 summary + add to canonical_index.json."""
import sys, io, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
TT_PATH = BASE / 'Tom-tat-tung-bai' / 'QT064_Stephens_Photovoice_ScopingReview_IntlJAdolYouth_2023.docx'

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)

# ===== Build summary docx =====
d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
DARK = RGBColor(31, 73, 125); RED = RGBColor(192, 0, 0); GREEN = RGBColor(54, 95, 44); ORANGE = RGBColor(191, 97, 14)

t = d.add_heading('QT064 — Tóm tắt: Photovoice trong NC sức khoẻ tâm thần VTN (Stephens 2023)', level=0)
for r in t.runs: r.font.color.rgb = DARK

meta = d.add_table(rows=8, cols=2); meta.style = 'Table Grid'
mrows = [
    ('Tên (VN)', 'Tổng quan phạm vi hệ thống về Photovoice trong nghiên cứu sức khoẻ tâm thần liên quan vị thành niên'),
    ('Tên (EN)', 'A systematic scoping review of Photovoice within mental health research involving adolescents'),
    ('Tác giả', 'Madison Stephens, Eleanor Keiller, Maev Conneely, Paul Heritage, Mariana Steffen, Victoria Jane Bird'),
    ('Đơn vị', 'WHO Collaborating Centre, Queen Mary University of London + UCL + East London NHS Foundation Trust, UK'),
    ('Tạp chí', 'International Journal of Adolescence and Youth, vol. 28, no. 1 (2023), Open Access'),
    ('DOI', '10.1080/02673843.2023.2244043'),
    ('Năm', '2023'),
    ('Mẫu + công cụ', '12 nghiên cứu Photovoice (~200 vị thành niên 10-22 tuổi); CASP + COREQ quality appraisal; PRISMA-ScR'),
]
for i, (k, v) in enumerate(mrows):
    c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
    meta.rows[i].cells[1].text = v
d.add_paragraph()

# Phương pháp
h = d.add_heading('Phương pháp', level=1)
for r in h.runs: r.font.color.rgb = DARK
d.add_paragraph('Scoping review (PRISMA-ScR). Tìm 5 databases: PSYCHinfo, PubMed, Scopus, Web of Science, CINAHAL. Tìm lại 06/2022 + 05/2023.')
d.add_paragraph('Tiêu chí: VTN 10-19 tuổi (định nghĩa WHO 2021); nghiên cứu nguyên bản dùng PHƯƠNG PHÁP PHOTOVOICE; mọi quốc gia, mọi ngôn ngữ; định tính (loại định lượng).')
d.add_paragraph('Phân tích: thematic analysis phản tư (Braun & Clarke 2006); chất lượng đánh giá bằng CASP + COREQ.')
d.add_paragraph('Lựa chọn nghiên cứu: 215 (Burn extraction) + 334 (cập nhật) → 7+3+2 grey lit = 12 nghiên cứu.')
d.add_paragraph()

# Kết quả
h = d.add_heading('Kết quả nghiên cứu', level=1)
for r in h.runs: r.font.color.rgb = DARK
for it in [
    'CHẤT LƯỢNG: KHÔNG có nghiên cứu nào "tốt"; 6/12 trung bình; 4/12 kém. Khu vực thiếu báo cáo: tuyển mộ, thu thập dữ liệu, reflexivity.',
    'PHÂN BỐ: 10/12 ở HICs (USA n=6, Canada n=3, UK n=1); 2/12 ở LMICs (Kenya, Nam Phi); 9/12 đô thị; nữ đại diện quá mức (127 vs 47 nam).',
    'CỠ MẪU: 4-58 người/NC; ≤10 ở 6 nghiên cứu — lệch khuyến nghị Wang & Burris 1997 (≤10).',
    'PHƯƠNG PHÁP: 7/12 dùng Photovoice ĐỘC LẬP; 3/12 kết hợp phỏng vấn/viết sáng tạo. 6/12 dùng KHUNG SHOWeD (See, Happening, Our lives, Why exist, Do).',
    'CAMERA: 4 dùng máy dùng-một-lần (28 ảnh); 2 dùng máy số; 1 dùng điện thoại; 5 không báo cáo.',
    'CHỦ ĐỀ CHÍNH (4 nhóm): (1) Đấu tranh → sức mạnh + thuộc về (coping/resilience/beliefs about self); (2) Gia đình + bạn bè (hy vọng + (mis)understanding); (3) An toàn + nghèo đói; (4) Điều trị + môi trường lâm sàng (autonomy adolescent).',
    'XU HƯỚNG: hầu hết NC từ 2021; lĩnh vực mới nổi 2008 trở đi.',
    'KHÁC HICs vs LMICs: LMICs đặt trong bối cảnh BỆNH TRUYỀN NHIỄM (HIV) + khu định cư không chính thức; tài trợ khác.',
]:
    d.add_paragraph(it, style='List Bullet')
d.add_paragraph()

# Điểm mạnh
h = d.add_heading('Điểm mạnh', level=1)
for r in h.runs: r.font.color.rgb = GREEN
for it in [
    'TỔNG QUAN ĐẦU TIÊN cùng loại — lấp khoảng trống evidence',
    'PRISMA-ScR + CASP + COREQ — methodology nghiêm ngặt cho scoping',
    'Open Access (CC BY 4.0) — tiếp cận rộng',
    '6 tác giả từ 3 đơn vị (Queen Mary, UCL, East London NHS) — perspective đa dạng',
]:
    d.add_paragraph('• ' + it)
d.add_paragraph()

# Hạn chế
h = d.add_heading('Hạn chế', level=1)
for r in h.runs: r.font.color.rgb = RED
for it in [
    '⚠ Bằng chứng yếu — scoping review không đo effect size, không thay thế MA',
    'Cỡ mẫu ~200 VTN qua 12 NC — quá nhỏ để khái quát',
    'Bias đa chiều: Tây phương + đô thị + nữ quá đại diện',
    'KHÔNG có NC nào "tốt" về chất lượng báo cáo',
    'Định nghĩa VTN có thể khác giữa các nguồn',
    'Pre-extracted data (Burn et al.) → khó nhân đôi',
]:
    d.add_paragraph('• ' + it)
d.add_paragraph()

# Phản biện
h = d.add_heading('Phản biện ngắn', level=1)
for r in h.runs: r.font.color.rgb = RED
p = d.add_paragraph()
r = p.add_run('Bài Photovoice ĐẦU TIÊN trong CSDL 87 canonical (chủ yếu RCT/SR/MA định lượng). Photovoice là phương pháp ĐỊNH TÍNH có sự tham gia, bổ sung cho phương pháp định lượng — KHÔNG thay thế. Đối với đề tài lo âu HS VN: bài này CÓ GIÁ TRỊ làm methodological reference, KHÔNG dùng cho effect size hay khuyến nghị can thiệp.'); r.font.color.rgb = RED
d.add_paragraph()

# Hướng nghiên cứu tiếp
h = d.add_heading('Hướng nghiên cứu tiếp theo', level=1)
for r in h.runs: r.font.color.rgb = DARK
for it in [
    'Áp dụng Photovoice cho HS THCS/THPT VN có triệu chứng lo âu (GAD-7 ≥ 5)',
    'Cấu trúc đề xuất: 30-40 HS, đào tạo 1 buổi → 1-2 tuần chụp ảnh chủ đề "điều làm tôi lo âu/an toàn ở trường" → 4 buổi thảo luận SHOWeD → triển lãm + advocacy',
    'Cân giới 50:50 (đề tài hiện thiếu nam tham gia định tính)',
    'Đa dạng vùng (đô thị Hà Nội/TPHCM + nông thôn Nghệ An/Lạng Sơn)',
    'Báo cáo theo COREQ ≥ 25 mục (đạt mức "tốt" — bài Stephens không có NC nào đạt!)',
    'Photovoice trực tuyến — tuyển ≥ 100 HS (Subasi 2023, Doyumğaç 2021)',
]:
    d.add_paragraph('• ' + it)
d.add_paragraph()

# Đánh giá
h = d.add_heading('Đánh giá chất lượng: ⭐⭐⭐⭐ (4/5)', level=1)
for r in h.runs: r.font.color.rgb = ORANGE
d.add_paragraph('Cao — scoping review đầu tiên về Photovoice trong sức khoẻ tâm thần VTN. Methodology PRISMA-ScR chuẩn, Open Access. Hạn chế ở cỡ mẫu nhỏ (12 NC), không đo effect size — đó là bản chất của scoping. Phù hợp làm methodological reference cho dự án Photovoice tại VN.')

d.add_paragraph()
p = d.add_paragraph()
r = p.add_run('Key fact cho RAG: '); r.bold = True; r.font.size = Pt(10)
p.add_run('Photovoice scoping review 12 NC, ~200 VTN. 4 chủ đề: coping/resilience, family/friends, safety/poverty, treatment. Bài Photovoice ĐẦU TIÊN trong CSDL — methodological reference cho thiết kế Photovoice VN.').font.size = Pt(10)

d.save(TT_PATH)
print(f'Summary saved: {TT_PATH.name} ({TT_PATH.stat().st_size//1024} KB)')

# ===== Update canonical_index.json =====
with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

idx['QT064'] = {
    'id': 'QT064',
    'descriptor': 'Stephens_Photovoice_ScopingReview_IntlJAdolYouth_2023',
    'summary': 'QT064_Stephens_Photovoice_ScopingReview_IntlJAdolYouth_2023.docx',
    'translation': 'QT064_Stephens_Photovoice_ScopingReview_IntlJAdolYouth_2023.docx',
    'pdf': 'QT064_Stephens_Photovoice_ScopingReview_IntlJAdolYouth_2023.pdf',
    'pdf_folder': 'The-gioi_Au-My-Uc',
    'doi': '10.1080/02673843.2023.2244043',
    'source_url': 'https://www.tandfonline.com/doi/full/10.1080/02673843.2023.2244043',
    'open_access_pdf_url': 'https://discovery.ucl.ac.uk/id/eprint/10179741/',
    'cao_date': '2026-04-29',
    'scope': 'main',
    'method': 'scoping_review_qualitative',
}

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

n_vn = sum(1 for k in idx if k.startswith('VN'))
n_qt = sum(1 for k in idx if k.startswith('QT'))
print(f'\nCanonical updated: {len(idx)} entries ({n_vn} VN + {n_qt} QT)')
print(f'  QT064 added with: pdf, summary, translation paths')
