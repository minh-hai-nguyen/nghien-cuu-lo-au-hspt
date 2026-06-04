# -*- coding: utf-8 -*-
"""Sample verify 10 QT suspicious de confirm pattern (false positive vs real error).
01/06/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import fitz

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# 10 QT suspicious - mix giua batch 1 + 2
samples = ['003', '005', '012', '021', '024', '028', '031', '034', '047', '063']


def find_pdf(qt_num):
    for root, dirs, files in os.walk(os.path.join(ROOT, '02_Papers-goc')):
        for f in files:
            if f.startswith(f'QT{qt_num}_') and f.lower().endswith('.pdf'):
                return os.path.join(root, f)
            if f.startswith(f'{int(qt_num):02d}_') and f.lower().endswith('.pdf'):
                return os.path.join(root, f)
    return None


def find_qt(qt_num):
    for f in os.listdir(os.path.join(ROOT, 'Tom-tat-tung-bai')):
        if f.startswith(f'QT{qt_num}_') and f.endswith('.docx') \
           and 'FIXED' not in f and 'BACKUP' not in f and 'BEFORE' not in f:
            return os.path.join(ROOT, 'Tom-tat-tung-bai', f)
    return None


for qt_num in samples:
    qt_p = find_qt(qt_num)
    pdf_p = find_pdf(qt_num)
    print(f'=== QT{qt_num} ===')
    if not qt_p:
        print(f'  NO QT file')
        continue
    if not pdf_p:
        print(f'  NO PDF available - cannot verify')
        continue

    # Read both
    try:
        d = Document(qt_p)
        qt_text = '\n'.join(p.text for p in d.paragraphs)
    except Exception as e:
        print(f'  QT read error: {e}'); continue

    try:
        doc = fitz.open(pdf_p)
        pdf_text = ''
        for i in range(doc.page_count):
            pdf_text += doc[i].get_text() + '\n'
        doc.close()
    except Exception as e:
        print(f'  PDF read error: {e}'); continue

    # Extract QT % claims
    qt_pcts = sorted(set(re.findall(r'\d+[,.]\d+\s*%', qt_text)))

    if not qt_pcts:
        print(f'  QT: 0 percentages - likely SR/MA/Narrative review → OK')
        # Check QT type from first paragraph
        first = next((p.text for p in d.paragraphs if p.text.strip()), '')
        print(f'    Type: {first[:200]}')
        continue

    # Check each
    matched_direct = 0
    matched_derived = 0
    real_miss = []

    for pct in qt_pcts:
        period = pct.replace(',', '.')
        compact = pct.replace(' ', '')
        period_compact = period.replace(' ', '')
        direct_find = (pct in pdf_text or period in pdf_text or
                       compact in pdf_text or period_compact in pdf_text)
        if direct_find:
            matched_direct += 1
        else:
            # Try if it could be sum of 2 nearby values
            # Get the number part
            try:
                val = float(period.replace('%', '').strip())
                # Check if val is reasonable for sum patterns
                real_miss.append((pct, val))
            except:
                real_miss.append((pct, None))

    total = len(qt_pcts)
    print(f'  {total} % claimed: {matched_direct} direct match, {len(real_miss)} miss')
    if real_miss[:5]:
        print(f'    Miss: {[m[0] for m in real_miss[:5]]}')

    # Quick verdict
    if matched_direct / total >= 0.5:
        print(f'  Verdict: OK (>=50% direct match)')
    elif matched_direct / total >= 0.25:
        print(f'  Verdict: PARTIAL - need manual verify')
    else:
        print(f'  Verdict: LOW match - manual verify needed')

    print()
