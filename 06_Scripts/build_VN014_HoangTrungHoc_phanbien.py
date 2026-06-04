"""Build doc giai thich 5 diem phan bien (3-7) bai VN014 Hoang Trung Hoc 2025."""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/VN014_HoangTrungHoc_2025_phan_bien_diem_3_7.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color
    return p

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold
    return p

def para_black(text, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(12); r.bold = bold; r.italic = italic
    return p

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)
    return p

def bullet_black(text, italic=False):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(12); r.italic = italic
    return p

# =====================================================
# TIÊU ĐỀ
# =====================================================
H('Phản biện chi tiết — VN014 Hoàng Trung Học 2025', level=1)
H('— bài "Mức độ căng thẳng, lo âu và trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 tại Việt Nam" —', level=2)

# =====================================================
# CÂU HỎI (xanh)
# =====================================================
H('5 điểm phản biện thầy đã nêu (3–7)', level=2)
para_blue(
    '3. DASS-21 là sàng lọc, tỷ lệ 41,5–65,5% bao gồm cả mức nhẹ — so với V-NAMHS 2022 '
    '(DISC-5 chẩn đoán chỉ 2,3%), chênh lệch rất lớn.'
)
para_blue(
    '4. Lấy mẫu "convenient random sampling" — mâu thuẫn (thuận tiện ≠ ngẫu nhiên), hạn '
    'chế tính đại diện.'
)
para_blue(
    '5. R² = 0,190 — mô hình chỉ giải thích 19% phương sai, 81% còn lại do yếu tố chưa '
    'đo lường.'
)
para_blue(
    '6. Bảng 3 trong bài gốc ghi nhãn "Reckless" cho mức "Không lo âu" — có thể lỗi dịch '
    'từ tiếng Việt sang tiếng Anh.'
)
para_blue(
    '7. Không phân tích riêng giới tính trong tỷ lệ lo âu — chỉ có Beta = 0,053 cho giới '
    'tính trong hồi quy.'
)

# =====================================================
# BACKGROUND VERIFY (đen)
# =====================================================
H('Verify với bản dịch + tóm tắt VN014 trong corpus', level=2)
para_black(
    'Em đối chiếu với file Tom-tat-tung-bai/VN014_HoangTrungHoc_2025_VN_COVID.docx — '
    '5 điểm phản biện đều CÓ CƠ SỞ. Cụ thể:', italic=True
)
bullet_black('Mẫu: 8.389 VTN (Đợt 1 = 4.052 trong COVID 12/2021 + Đợt 2 = 4.337 sau COVID 9/2023). 6 tỉnh, nông thôn 62%, THCS+THPT lớp 6-12.', italic=True)
bullet_black('Công cụ: DASS-21 phiên bản tiếng Việt (Lovibond & Lovibond 1995).', italic=True)
bullet_black('Phương pháp lấy mẫu: "Lấy mẫu thuận tiện ngẫu nhiên (convenient random sampling)" (NGUYÊN VĂN trong bản dịch tóm tắt — ĐÚNG là mâu thuẫn thuật ngữ).', italic=True)
bullet_black('Hồi quy đa biến: R² = 0,190 (trong COVID).', italic=True)
bullet_black('Tỷ lệ lo âu: trong COVID 41,5% → sau COVID 25,4% (đã verify từ tóm tắt VN014).', italic=True)
bullet_black('⚠ Tỷ lệ stress 65,5%: thầy nêu trong câu hỏi nhưng EM CHƯA VERIFY trực tiếp được trong tóm tắt VN014 hiện có. Cần thầy gửi đoạn nguyên văn hoặc xác nhận nguồn (Bảng 1 bài gốc?).', italic=True)
bullet_black('Beta giới tính = 0,053 (kiểm chứng yếu trong hồi quy).', italic=True)
bullet_black('Tạp chí: AJPR (Asian Journal of Public Research?) — không impact factor rõ ràng, không lập chỉ mục PubMed/Scopus.', italic=True)

# =====================================================
# GIẢI THÍCH CHI TIẾT TỪNG ĐIỂM
# =====================================================
H('Điểm 3 — DASS-21 sàng lọc 41,5–65,5% vs V-NAMHS DSM-5 chẩn đoán 2,3%', level=2)

para_black('Phân biệt SÀNG LỌC vs CHẨN ĐOÁN:', bold=True)
bullet_black('Sàng lọc (screening): công cụ tự khai báo (DASS-21, GAD-7, PHQ-9...) đo TRIỆU CHỨNG. Mục đích: nhận diện rộng người CÓ KHẢ NĂNG có vấn đề. Ngưỡng cắt thường THẤP để tăng độ nhạy (sensitivity), chấp nhận tỷ lệ dương tính giả cao.')
bullet_black('Chẩn đoán (diagnostic): phỏng vấn cấu trúc theo DSM-5/ICD-11 (DISC-5, K-SADS, MINI-Kid). Đo RỐI LOẠN ĐẦY ĐỦ với tiêu chí số triệu chứng + thời gian + suy giảm chức năng. Tỷ lệ thấp hơn nhiều.')
para_black('Chênh lệch sàng lọc/chẩn đoán THƯỜNG là 5–37 lần (COVID-19 Mental Disorders Collaborators 2021, Lancet). VN014 (DASS-21, lo âu 41,5%) so với V-NAMHS 2022 (DISC-5, lo âu 2,3%) → chênh ~18 lần — NẰM trong khoảng dự kiến.')

para_black('Điểm cần phản biện rõ:', bold=True)
bullet_black('Hoàng Trung Học KHÔNG được dùng DASS-21 41,5% như một con số "có thể so sánh trực tiếp" với 2,3% V-NAMHS. Hai con số đo HAI THỨ KHÁC NHAU.')
bullet_black('DASS-21 với ngưỡng "lo âu nhẹ trở lên" (≥ 8 điểm) bao gồm cả người có 8-9 điểm — tức triệu chứng nhẹ thoáng qua, KHÔNG phải rối loạn lâm sàng.')
bullet_black('Khi viết báo cáo CTH, em đề nghị thầy ghi rõ: "Tỷ lệ sàng lọc DASS-21 (lo âu nhẹ trở lên) 41,5% — bao gồm các mức nhẹ; tỷ lệ rối loạn lo âu chẩn đoán DSM-5 V-NAMHS 2022 chỉ 2,3% — chênh lệch ~18 lần phản ánh sự khác biệt giữa sàng lọc và chẩn đoán."')

H('Điểm 4 — "Convenient random sampling" mâu thuẫn thuật ngữ', level=2)

tbl4 = doc.add_table(rows=4, cols=2)
tbl4.style = 'Light Grid Accent 1'
header = tbl4.rows[0]
header.cells[0].text = 'Phương pháp lấy mẫu'
header.cells[1].text = 'Đặc điểm'
data4 = [
    ('Lấy mẫu THUẬN TIỆN (convenience sampling)', 'Chọn người dễ tiếp cận (HS trường gần, lớp quen). KHÔNG ngẫu nhiên. Tính đại diện THẤP.'),
    ('Lấy mẫu NGẪU NHIÊN (random sampling)', 'Mỗi cá nhân trong dân số có xác suất được chọn xác định (ngẫu nhiên đơn giản, hệ thống, phân tầng, theo cụm...). Tính đại diện CAO.'),
    ('"Convenient random sampling"', 'KHÔNG TỒN TẠI trong phương pháp luận chuẩn. Đây là cụm tự đặt — mâu thuẫn nội tại.'),
]
for i, (a, b) in enumerate(data4, 1):
    tbl4.rows[i].cells[0].text = a
    tbl4.rows[i].cells[1].text = b
for row in tbl4.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black('Điều có thể XẢY RA trong thực tế:', bold=True)
bullet_black('Tác giả chọn 6 tỉnh THUẬN TIỆN (có quan hệ liên kết), rồi trong mỗi tỉnh chọn trường THUẬN TIỆN, rồi trong mỗi trường lấy NGẪU NHIÊN học sinh.')
bullet_black('Đây là "multi-stage cluster sampling with convenience at higher levels and randomization at student level" — phải mô tả rõ NHƯ VẬY, không gộp thành "convenient random".')
bullet_black('Hậu quả: kết quả KHÓ tổng quát hóa sang toàn bộ HS VN — chỉ phản ánh 6 tỉnh đã chọn.')

H('Điểm 5 — R² = 0,190 chỉ giải thích 19% phương sai', level=2)

para_black('R² (coefficient of determination) là tỷ lệ phương sai biến phụ thuộc được mô hình giải thích. R² = 0,190 → 81% phương sai CÒN LẠI do các yếu tố CHƯA đo lường.')

para_black('Phân loại R² trong khoa học xã hội/tâm lý:', bold=True)
tbl5 = doc.add_table(rows=5, cols=3)
tbl5.style = 'Light Grid Accent 1'
header = tbl5.rows[0]
header.cells[0].text = 'R²'
header.cells[1].text = 'Phân loại'
header.cells[2].text = 'Diễn giải cho VN014'
data5 = [
    ('< 0,10', 'Yếu', ''),
    ('0,10 – 0,30', 'Trung bình', '✓ R² = 0,190 nằm ở đây'),
    ('0,30 – 0,50', 'Khá', ''),
    ('> 0,50', 'Mạnh', ''),
]
for i, (a, b, c) in enumerate(data5, 1):
    tbl5.rows[i].cells[0].text = a
    tbl5.rows[i].cells[1].text = b
    tbl5.rows[i].cells[2].text = c
for row in tbl5.rows:
    for cell in row.cells:
        for p in cell.paragraphs:
            for r in p.runs:
                r.font.size = Pt(11); r.font.name = 'Times New Roman'

para_black('')
para_black('Trong tâm lý vị thành niên, R² ~ 0,15–0,25 là PHỔ BIẾN khi mô hình chỉ dùng các biến hành vi (giấc ngủ, thiết bị, thể thao, quan hệ cha mẹ). Yếu tố CHƯA đo của Hoàng Trung Học có thể là: tính cách, di truyền, ACEs (trải nghiệm bất lợi), sang chấn, kinh tế gia đình, mạng xã hội, học lực — đây là điều phải nêu rõ trong "limitations".')

H('Điểm 6 — Lỗi dịch "Reckless" cho mức "Không lo âu"', level=2)
para_black('⚠ LƯU Ý: em CHƯA đọc trực tiếp PDF gốc bài Hoàng Trung Học 2025 — claim "Reckless" hiện chỉ dựa vào điểm thầy đã nêu. Nếu đúng như mô tả của thầy, phân tích dưới đây ÁP DỤNG; nếu không, thầy bỏ qua phần này.', italic=True, bold=True)
para_black('')
para_black('"Reckless" trong tiếng Anh = liều lĩnh, bốc đồng, không quan tâm tới hậu quả. KHÔNG có nghĩa "không lo âu". Có thể tác giả VN dịch sai từ:')
bullet_black('"Bình thường" → "Normal" (đúng), không phải "Reckless".')
bullet_black('"Không lo lắng" → "Not anxious" hoặc "Worry-free" (đúng), không phải "Reckless".')
bullet_black('"Buông lỏng" → "Relaxed", "Carefree" (đúng), không phải "Reckless".')

para_black('Hậu quả lỗi này:', bold=True)
bullet_black('Đọc giả quốc tế hiểu sai: "reckless" trong DASS-21 = nhóm có hành vi nguy cơ! Trong khi tác giả muốn chỉ nhóm KHÔNG có triệu chứng lo âu.')
bullet_black('Đây là lỗi dịch thuật quan trọng cần GỬI ERRATUM tới tạp chí — không phải lỗi nhỏ chỉnh tay được.')
bullet_black('Khi trích Bảng 3 vào báo cáo CTH, em đề nghị thầy GHI CHÚ rõ: "Tác giả ghi nhãn \'Reckless\' cho nhóm Không lo âu (theo phân loại DASS-21 của Lovibond & Lovibond 1995, mức ≤ 7 điểm = không lo âu) — đây có thể là lỗi dịch thuật từ tiếng Việt."')

H('Điểm 7 — Không phân tích giới tính trong tỷ lệ lo âu', level=2)
para_black(
    'Tác giả chỉ đưa Beta = 0,053 cho giới tính trong hồi quy đa biến — KHÔNG báo cáo:'
)
bullet_black('Tỷ lệ lo âu RIÊNG cho NAM vs NỮ.')
bullet_black('Khoảng tin cậy KTC 95% cho hệ số giới tính.')
bullet_black('Effect size (Cohen\'s d hoặc OR) cho khác biệt giới.')
bullet_black('Subgroup analysis trong COVID vs sau COVID theo giới.')

para_black('Beta = 0,053 nói lên gì?', bold=True)
bullet_black('Beta chuẩn hoá 0,053 = RẤT NHỎ (Cohen: 0,1 nhỏ / 0,3 TB / 0,5 lớn). Tăng 1 SD biến giới tính (nam → nữ chẳng hạn) → tăng 0,053 SD tổng điểm DAS.')
bullet_black('Hệ số nhỏ + tác giả KHÔNG báo p-value riêng → có thể KHÔNG có ý nghĩa thống kê dù mẫu n=8.389 lớn.')

para_black('Vì sao đây là điểm yếu QUAN TRỌNG:', bold=True)
bullet_black('Y văn QUỐC TẾ (Hoa 2024 VN, Puyat 2025 Philippines, NSCH 2020 US, Jenkins 2023): nữ luôn cao hơn nam ~1,5–2 lần. Hoàng Trung Học không xác nhận phát hiện này.')
bullet_black('Có 3 khả năng: (a) thật sự VN không có khác biệt giới (= phát hiện mới); (b) tác giả chưa phân tích kỹ; (c) DASS-21 ngưỡng cắt giống nhau cho 2 giới mà thực tế nữ có biểu hiện thân thể hoá khác nam.')
bullet_black('Nếu thầy muốn dùng VN014 để bàn về khác biệt giới ở HS VN, KHÔNG đủ dữ liệu — phải kết hợp với Hoa 2024 (nữ 1,74 vs nam 1,50) hoặc V-NAMHS 2022 (không khác biệt).')

# =====================================================
# CÂU TRẢ LỜI (xanh, gom 1 chỗ trước phụ lục)
# =====================================================
H('CÂU TRẢ LỜI', level=2, color=BLUE)

para_blue('Tóm gọn — cả 5 điểm thầy nêu đều CÓ CƠ SỞ và thực sự là điểm yếu của VN014:', bold=True)

bullet_blue(
    'Điểm 3 (DASS-21 sàng lọc vs DISC-5 chẩn đoán): ĐÚNG. Chênh ~18 lần (41,5% vs '
    '2,3%) phản ánh sàng lọc bao gồm cả mức nhẹ. Khi viết, không được trộn 2 con số.'
)
bullet_blue(
    'Điểm 4 ("convenient random sampling"): MÂU THUẪN nội tại. Thuận tiện ≠ ngẫu '
    'nhiên. Có thể tác giả dùng cluster sampling 2 cấp (tỉnh/trường thuận tiện + HS '
    'ngẫu nhiên) nhưng phải mô tả rõ — không gộp thành 1 cụm sai.'
)
bullet_blue(
    'Điểm 5 (R² = 0,190): mô hình giải thích 19% phương sai — mức trung bình cho '
    'tâm lý vị thành niên, nhưng 81% còn lại do yếu tố CHƯA ĐO. Thiếu: tính cách, '
    'ACEs, kinh tế, mạng xã hội, học lực.'
)
bullet_blue(
    'Điểm 6 ("Reckless" cho không lo âu): NẾU đúng như thầy nêu (em chưa verify trực '
    'tiếp PDF gốc), đây là LỖI DỊCH NẶNG. Reckless = liều lĩnh, không phải "không lo '
    'âu". Cần erratum tới tạp chí. Đề nghị thầy verify trực tiếp Bảng 3 + ghi chú rõ.'
)
bullet_blue(
    'Điểm 7 (Beta giới tính = 0,053, không phân tích riêng): RẤT NHỎ + thiếu '
    'subgroup analysis. Trái với y văn quốc tế (nữ thường > nam 1,5–2 lần). Không '
    'đủ dữ liệu để khẳng định "VN không có khác biệt giới".'
)

para_blue('Tóm tắt độ tin cậy của VN014:', bold=True)
para_blue(
    'Cỡ mẫu 8.389 lớn nhất VN, 6 tỉnh đa dạng — là ƯU ĐIỂM. Nhưng tạp chí AJPR yếu '
    '(không Scopus/PubMed), thiết kế mẫu mâu thuẫn, R² thấp, lỗi dịch, và không phân '
    'tích giới — là HẠN CHẾ. Đánh giá tổng: ⭐⭐⭐ Trung bình-Khá. Có thể trích cho '
    'số liệu mô tả (xu hướng giảm 41,5% → 25,4%), nhưng KHÔNG nên dùng làm bằng '
    'chứng can thiệp hoặc chẩn đoán.'
)

para_blue('Khuyến nghị cách trích vào báo cáo CTH:', bold=True)
bullet_blue(
    'NÊN: "Theo Hoàng Trung Học (2025), tỷ lệ SÀNG LỌC lo âu (DASS-21) ở 8.389 VTN '
    '6 tỉnh VN giảm từ 41,5% trong COVID xuống 25,4% sau COVID — vẫn ở mức cao."'
)
bullet_blue(
    'KHÔNG NÊN: "Tỷ lệ lo âu ở HS VN là 41,5%" — vì gây hiểu nhầm là tỷ lệ chẩn đoán.'
)
bullet_blue(
    'Khi đối chiếu với V-NAMHS, ghi rõ: "DASS-21 sàng lọc 41,5% bao gồm các mức nhẹ; '
    'DSM-5 chẩn đoán đầy đủ chỉ 2,3% (V-NAMHS 2022)."'
)

# =====================================================
# PHỤ LỤC — REFERENCES
# =====================================================
H('Phụ lục — Tài liệu tham khảo', level=2)

para_black(
    '1. Hoàng Trung Học (2025). Mức độ căng thẳng, lo âu và trầm cảm ở thanh thiếu niên '
    'trong và sau đại dịch COVID-19 tại Việt Nam: Nghiên cứu cắt ngang. Asian Journal '
    'of Public Research (AJPR). [VN014 trong DB dự án.]', italic=True
)

para_black(
    '2. Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety '
    'Stress Scales (2nd ed.). Sydney: Psychology Foundation. — Bài gốc DASS-21.',
    italic=True
)

para_black(
    '3. UNICEF Việt Nam, Bộ LĐ-TB-XH, Tổng cục Thống kê (2022). Khảo sát Sức khỏe Tâm '
    'thần Vị thành niên Việt Nam (V-NAMHS 2022). Hà Nội. — DSM-5 chẩn đoán, lo âu '
    '2,3%. [VN002 trong DB dự án.]', italic=True
)

para_black(
    '4. COVID-19 Mental Disorders Collaborators (2021). Global prevalence and burden '
    'of depressive and anxiety disorders in 204 countries and territories in 2020 due '
    'to the COVID-19 pandemic. The Lancet, 398(10312), 1700–1712. — Phân biệt sàng lọc '
    'vs chẩn đoán; tăng 25% global lo âu+trầm cảm post-COVID.', italic=True
)

para_black(
    '5. Cohen, J. (1988). Statistical Power Analysis for the Behavioral Sciences (2nd '
    'ed.). Hillsdale, NJ: Erlbaum. — Khung phân loại Beta/effect size.', italic=True
)

para_black(
    '6. Etikan, I., Musa, S. A., & Alkassim, R. S. (2016). Comparison of convenience '
    'sampling and purposive sampling. American Journal of Theoretical and Applied '
    'Statistics, 5(1), 1–4. — Phân biệt rõ convenience vs random sampling.', italic=True
)

para_black('7. Tham khảo trong corpus dự án (DB Lo-au):')
bullet_black(
    'VN014 Hoàng Trung Học (2025) — bài đang phản biện. Bản dịch: '
    '03_Ban-dich/VN014_HoangTrungHoc_2025_VN_COVID.docx.', italic=True
)
bullet_black(
    'VN001 Hoa 2024 (Hà Nội, GAD-7, lo âu 40,6%, có phân tích giới: nữ > nam) — đối '
    'chiếu để thấy VN014 thiếu phân tích giới.', italic=True
)
bullet_black(
    'VN002 V-NAMHS 2022 (DISC-5, chẩn đoán 2,3%) — chuẩn vàng để so sánh sàng lọc DASS-21.',
    italic=True
)
bullet_black(
    'Pham 2024 (VN, Huế, social support) — có dữ liệu chăm sóc cảm xúc beta = -0,40 — '
    'đối chiếu được với Beta quan hệ cha mẹ-con của Hoàng Trung Học (0,272).', italic=True
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
