# -*- coding: utf-8 -*-
"""Vong 3 - Deep edge cases.
- Vong 11: Word choice + terminology consistency
- Vong 12: Statistical reporting style
- Vong 13: Brackets + punctuation integrity
- Vong 14: Abbreviation expansion check
- Vong 15: Final cross-validation vs LA Kết luận
29/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', 'TrichYeuLA_CongThiHang_v2_29052026.docx')
LA = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

d = Document(FILE)
la = Document(LA)

all_d = '\n'.join(p.text for p in d.paragraphs)
for tb in d.tables:
    for row in tb.rows:
        for c in row.cells:
            all_d += '\n' + c.text

all_la = '\n'.join(p.text for p in la.paragraphs)
for tb in la.tables:
    for row in tb.rows:
        for c in row.cells:
            all_la += '\n' + c.text

en_start_idx = None
for i, p in enumerate(d.paragraphs):
    if 'DISSERTATION ABSTRACT' in p.text:
        en_start_idx = i
        break
vn_text = '\n'.join(p.text for p in d.paragraphs[:en_start_idx])
en_text = '\n'.join(p.text for p in d.paragraphs[en_start_idx:])

issues = []
total = 0
passed = 0


def check(name, cond, detail=''):
    global total, passed
    total += 1
    status = '✓' if cond else '✗'
    print(f'  {status} {name}')
    if detail:
        print(f'      {detail}')
    if cond:
        passed += 1
    else:
        issues.append(name)
    return cond


# ============================================================
# VONG 11: WORD CHOICE + TERMINOLOGY CONSISTENCY
# ============================================================
print('=== VÒNG 11: WORD CHOICE + TERMINOLOGY ===')

# "rối loạn lo âu" preferred over standalone "lo âu" — count ratio
rlla_cnt = vn_text.lower().count('rối loạn lo âu')
loau_only_cnt = vn_text.lower().count('lo âu') - rlla_cnt
check('11.1 "rối loạn lo âu" sử dụng đủ chiếm ưu thế (>=10 lần)',
      rlla_cnt >= 10, f'RLLA={rlla_cnt}, lo âu standalone~{loau_only_cnt}')

# Anxiety disorders (EN)
ad_cnt = en_text.lower().count('anxiety disorder')
check('11.2 "anxiety disorder" (EN) >=8 lần',
      ad_cnt >= 8, f'count={ad_cnt}')

# Consistent: "học sinh trung học cơ sở" not "học sinh THCS" in main body
check('11.3 "trung học cơ sở" >= 5 lần (full form)',
      vn_text.lower().count('trung học cơ sở') >= 5)

# Brief COPE 3 strategies mentioned (giải quyết / hỗ trợ / né tránh)
check('11.4 3 chiến lược đối phó: giải quyết vấn đề + hỗ trợ + né tránh',
      'giải quyết vấn đề' in vn_text.lower()
      and 'hỗ trợ' in vn_text.lower()
      and ('né tránh' in vn_text.lower() or 'tránh né' in vn_text.lower()))

check('11.5 EN: problem-solving + social support',
      'problem-solving' in en_text.lower() and 'social support' in en_text.lower())

# Check term "phương trình cấu trúc" (SEM expansion)
check('11.6 "phương trình cấu trúc" (SEM expansion VN)',
      'phương trình cấu trúc' in vn_text.lower())

check('11.7 "structural equation model" (SEM expansion EN)',
      'structural equation model' in en_text.lower())

# Validate "ĐTB" Vietnamese vs "M" English
check('11.8 VN dùng "ĐTB" (Điểm trung bình)',
      'ĐTB' in vn_text)
check('11.9 EN dùng "M =" (Mean)',
      'M =' in en_text)


# ============================================================
# VONG 12: STATISTICAL REPORTING STYLE
# ============================================================
print('\n=== VÒNG 12: STATISTICAL REPORTING STYLE ===')

# Italic conventions checked manually — APA wants F italic
# Here just check value formatting

# F values present
check('12.1 F values reported with proper format',
      'F = 0,246' in vn_text and 'F = 0.246' in en_text)

# p values report
check('12.2 p values với so sánh "<" hoặc "="',
      ('p < 0,001' in vn_text or 'p = 0,001' in vn_text)
      and ('p < 0.001' in en_text or 'p = 0.001' in en_text))

# Effect size language
check('12.3 Cường độ tác động (effect size) mentioned',
      'cường độ' in vn_text.lower() and 'magnitude' in en_text.lower())

# β coefficient or descriptor (not strict — mentioned in body but may not appear in trích yếu)
check('12.4 Risk factor ordering with > separator',
      '>' in vn_text and '>' in en_text)

# Number ranges 85-89 with en-dash
check('12.5 VN: "85–89" với en-dash hoặc "85-89"',
      ('85–89' in vn_text or '85-89' in vn_text))
check('12.6 EN: "85–89" với en-dash hoặc "85-89"',
      ('85–89' in en_text or '85-89' in en_text))

# CI brackets if mentioned (not strict)
# Percentages: should use %
check('12.7 % sign used for percentages',
      '%' in all_d)

# No raw "p value" - should be "p" italic in APA, but just check formatting
check('12.8 "p" với spaces "p < 0," (proper spacing)',
      ' p < 0,' in vn_text or 'p < 0,001' in vn_text)


# ============================================================
# VONG 13: BRACKETS + PUNCTUATION INTEGRITY
# ============================================================
print('\n=== VÒNG 13: BRACKETS + PUNCTUATION ===')

# Balanced parentheses
def balanced(text, open_c, close_c):
    return text.count(open_c) == text.count(close_c)

check('13.1 Parentheses cân bằng ( )',
      balanced(all_d, '(', ')'),
      f'open={all_d.count("(")} close={all_d.count(")")}')

check('13.2 Square brackets cân bằng [ ]',
      balanced(all_d, '[', ']'),
      f'open={all_d.count("[")} close={all_d.count("]")}')

check('13.3 Quotation marks " " cân bằng',
      all_d.count('"') % 2 == 0,
      f'count={all_d.count(chr(34))}')

# No double punctuation
check('13.4 KHÔNG có double punctuation "..."',
      '...' not in all_d.replace('Anxiety...', '').replace('NCS điền', ''))
check('13.5 KHÔNG có ",,"', ',,' not in all_d)
check('13.6 KHÔNG có ";;" ', ';;' not in all_d)
check('13.7 KHÔNG có ":,"', ':,' not in all_d)
check('13.8 KHÔNG có ",."', ',.' not in all_d)
# 13.9: ".," = chuẩn APA "et al., 2011" → exclude pattern này khi check
text_excl_et_al = re.sub(r'et al\.,', 'et al', all_d)
check('13.9 KHÔNG có ".," (loại trừ "et al.," chuẩn APA)', '.,' not in text_excl_et_al)

# No leading/trailing whitespace artifacts
check('13.10 KHÔNG có " ;"', ' ;' not in all_d)
check('13.11 KHÔNG có " ."', not bool(re.search(r' \.[^a-z0-9]', all_d)))
check('13.12 KHÔNG có "  " (2 spaces) lỗi',
      not bool(re.search(r'[a-zA-Zàáảãạăắằẳẵặâấầẩẫậ]  [a-zA-Z]', all_d)))


# ============================================================
# VONG 14: ABBREVIATION EXPANSION CHECK
# ============================================================
print('\n=== VÒNG 14: ABBREVIATION EXPANSION ===')

# DSM-5 used at least once with what it stands for? Not required in trích yếu
# But scales should have acronym + full name + author
acronyms_full = {
    'RCADS': 'Revised Children',
    'ESSA': 'Educational Stress Scale',
    'SAS-SV': 'Smartphone Addiction',
    'OBVQ': 'Olweus',
    'PSSM': 'Psychological Sense of School Membership',
    'MSPSS': 'Multidimensional Scale of Perceived Social Support',
    'RSES': 'Rosenberg',
    'Brief COPE': 'Carver',
}

for acr, full_part in acronyms_full.items():
    cnt_d = all_d.count(acr)
    has_full = full_part in all_d
    check(f'14.{acr} có cả acronym + full name "{full_part}..."',
          cnt_d >= 2 and has_full)

# CFA + SEM expanded
check('14.99a CFA + "nhân tố khẳng định" (VN expansion)',
      'CFA' in vn_text and 'nhân tố khẳng định' in vn_text.lower())
check('14.99b CFA + "confirmatory factor analysis" (EN expansion)',
      'CFA' in en_text and 'confirmatory factor analysis' in en_text.lower())


# ============================================================
# VONG 15: FINAL CROSS-VALIDATION vs LA KẾT LUẬN
# ============================================================
print('\n=== VÒNG 15: CROSS-VALIDATION vs LA KẾT LUẬN ===')

# These claims MUST exist in LA kết luận section
la_claims = [
    ('Khung chương trình tập huấn phòng ngừa rối loạn lo âu',
     'Khung CT đã đề xuất trong LA'),
    ('cường độ tác động tương đương (~85–89%)',
     'TTr 85-89%'),
    ('áp lực học tập (ĐTB = 51,13) > nghiện điện thoại (28,38) > bắt nạt thể chất (13,52)',
     'Thứ tự YTNC + ĐTB'),
    ('F = 0,246; p = 0,620',
     'F + p cho lo âu chia ly'),
    ('8 nội dung cơ bản',
     'Khung CT 8 nội dung'),
]

for claim, name in la_claims:
    has = claim in all_la
    check(f'15. LA chứa: "{name}"', has)

# Verify trích yếu derives directly from LA Kết luận content
key_phrases_la = [
    'lan tỏa',
    'xã hội',
    'chia ly',
    'áp lực học tập',
    'nghiện điện thoại',
    'bắt nạt',
    'lòng tự trọng',
    'gắn bó với trường học',
    'hỗ trợ',
    'đối phó',
]
for phrase in key_phrases_la:
    has_d = phrase in all_d.lower()
    has_la = phrase in all_la.lower()
    check(f'15.b "{phrase}" có ở cả Trích yếu + LA',
          has_d and has_la)


# ============================================================
# RESULT
# ============================================================
print()
print('=' * 70)
print(f'TỔNG VÒNG 11-15: {passed}/{total} checks PASS ({len(issues)} lỗi)')
if issues:
    print('\nDANH SÁCH LỖI:')
    for i, iss in enumerate(issues, 1):
        print(f'  {i}. {iss}')
    sys.exit(1)
else:
    print('\nSẠCH LỖI VÒNG 3!')
    sys.exit(0)
