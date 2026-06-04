# -*- coding: utf-8 -*-
"""
Sua loi VN022 bi propagate xuong cac downstream files:
- 249 GV -> 66 GV
- 47% hoc them -> 50% hoc them (EN goc ro rang: "50% of students report more than 3 hours per week")
- Them nhat quan ve 4 tinh thuc te (Gia Lai, khong phai Kon Tum/Ninh Thuan)

Target files:
- Bao cao v5-5 (latest, khong sua v3-v4b la snapshots cu)
- Tong hop lien bai 10042026 (latest)
- Tom tat QT037, QT038, QT042, QT046, QT047, QT051
- Memory files (neu co)
"""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Files to fix (only LATEST versions)
TARGETS = [
    os.path.join(ROOT, '01_Bao-cao',
                 'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-5.docx'),
    os.path.join(ROOT, '01_Bao-cao',
                 'Tổng hợp liên bài báo - Lo âu HS - 10042026.docx'),
    os.path.join(ROOT, 'Tom-tat-tung-bai',
                 'QT037_Praptomojati_CA-CBT_SEA_AJP_2024.docx'),
    os.path.join(ROOT, 'Tom-tat-tung-bai',
                 'QT038_DeSilva_SriLanka_RCT_2024.docx'),
    os.path.join(ROOT, 'Tom-tat-tung-bai',
                 'QT042_BrownCarter_UK_School_JMH_2025.docx'),
    os.path.join(ROOT, 'Tom-tat-tung-bai',
                 'QT046_Jagiello_AcademicStress_SR_2025.docx'),
    os.path.join(ROOT, 'Tom-tat-tung-bai',
                 'QT047_Dong_PLOS_DASS_YaAn_2025.docx'),
    os.path.join(ROOT, 'Tom-tat-tung-bai',
                 'QT051_Menon_LMIC_SEA_Pacific_APJPH_2025_ABSTRACT.docx'),
]

FIXES = [
    # Additional patterns with space/no-space variants
    ('47 % HS học thêm > 3h/tuần', '50 % HS học thêm > 3 h/tuần, 15 % > 9 h/tuần', '47->50 w space'),
    ('47 % HS học thêm >3h/tuần', '50 % HS học thêm >3 h/tuần, 15 % > 9 h/tuần', '47->50 w space alt'),
    ('47% học thêm > 3h/tuần', '50% học thêm > 3 h/tuần, 15% > 9 h/tuần', '47->50 variant'),
    ('(47% học thêm', '(50% học thêm', '47 in parens'),
    ('(47 % học thêm', '(50 % học thêm', '47 space parens'),
    ('47% HS học thêm', '50% HS học thêm', '47->50 HS no space'),
    ('47 % HS học thêm', '50 % HS học thêm', '47->50 HS w space'),
    # (find, replace, note)
    ('668 HS + 249 GV', '668 HS + 66 GV', 'Teacher count 249 -> 66 (real UNICEF 2022)'),
    ('249 GV | SDQ', '66 GV | SDQ', 'Table teacher count'),
    ('668 HS + 249 giáo viên', '668 HS + 66 giáo viên', 'VN Long form'),
    ('47 % HS Việt Nam học thêm > 3 giờ/tuần',
     '50 % HS Việt Nam học thêm > 3 giờ/tuần (28 % học sinh học > 3 giờ/đêm; 15 % > 9 giờ/tuần)',
     '47% -> 50% extra class'),
    ('47 % học sinh học thêm trên 3 giờ mỗi tuần',
     '50 % học sinh học thêm trên 3 giờ mỗi tuần (và 15 % học sinh học thêm trên 9 giờ mỗi tuần)',
     '47% -> 50% extra class'),
    ('47 % học thêm > 3 h/tuần', '50 % học thêm > 3 h/tuần', '47->50 short'),
    ('47% học thêm >3h/tuần', '50% học thêm >3h/tuần', '47->50 short no space'),
    ('47% học thêm >3h; bắt nạt',
     '50% học thêm >3h/tuần (15% >9h); bắt nạt',
     '47->50 + add 15% detail'),
    ('47% HS học thêm >3h/tuần', '50% HS học thêm >3h/tuần', '47->50 HS variant'),
    ('UNICEF VN22 (47% HS học thêm >3h/tuần)',
     'UNICEF VN22 (50% HS học thêm >3h/tuần, 15% > 9h/tuần)',
     'Full context'),
    ('UNICEF VN22 (47% học thêm >3h/tuần)',
     'UNICEF VN22 (50% học thêm >3h/tuần, 15% > 9h/tuần)',
     'Full context'),
    ('5 tỉnh VN | 668 HS + 249 GV',
     '4 tỉnh VN thực tế (HN, ĐB, GL, ĐT) + TPHCM dropped do COVID | 668 HS + 66 GV',
     'Provinces + teacher count'),
    ('VN022 | UNICEF VN 2022 | 47 % học thêm > 3 h/tuần',
     'VN022 | UNICEF VN 2022 | 50 % HS học thêm > 3 h/tuần, 15 % > 9 h/tuần',
     'Table row fix'),
]

total_changes = 0
for path in TARGETS:
    if not os.path.exists(path):
        print(f'SKIP (not found): {os.path.basename(path)}')
        continue
    try:
        d = Document(path)
        changes = 0
        # Fix paragraphs
        for p in d.paragraphs:
            if p.text.strip():
                new_text = p.text
                for find, replace, note in FIXES:
                    if find in new_text:
                        new_text = new_text.replace(find, replace)
                if new_text != p.text:
                    # Update by clearing all runs and setting first run
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = new_text
                    changes += 1
        # Fix tables
        for t in d.tables:
            for row in t.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        if p.text.strip():
                            new_text = p.text
                            for find, replace, note in FIXES:
                                if find in new_text:
                                    new_text = new_text.replace(find, replace)
                            if new_text != p.text:
                                for run in p.runs:
                                    run.text = ''
                                if p.runs:
                                    p.runs[0].text = new_text
                                changes += 1
        if changes > 0:
            d.save(path)
            print(f'✓ Fixed {os.path.basename(path)}: {changes} changes')
        else:
            print(f'  No changes needed: {os.path.basename(path)}')
        total_changes += changes
    except Exception as e:
        print(f'ERR {os.path.basename(path)}: {e}')

print(f'\nTOTAL CHANGES: {total_changes}')
