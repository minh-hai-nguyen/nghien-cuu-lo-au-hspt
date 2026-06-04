# -*- coding: utf-8 -*-
"""QT33 JAMA Screen (3t+3h) + QT34 Korea (4t+1h)"""
import sys,io
try:sys.stdout=io.TextIOWrapper(sys.stdout.buffer,encoding='utf-8')
except:pass
from docx import Document
from docx.shared import Pt,Cm,RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
RED=RGBColor(0xFF,0,0);BLUE=RGBColor(0,0x70,0xC0)
def md():
    d=Document();s=d.styles['Normal'];s.font.name='Times New Roman';s.font.size=Pt(12)
    s.paragraph_format.space_after=Pt(4);s.paragraph_format.line_spacing=1.5
    for sec in d.sections:sec.top_margin=Cm(2.5);sec.bottom_margin=Cm(2.5);sec.left_margin=Cm(3);sec.right_margin=Cm(2)
    return d
def rb(d,t):p=d.add_paragraph();r=p.add_run(t);r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=RED
def bl(d,t,bold=False):p=d.add_paragraph();r=p.add_run(t);r.font.name='Times New Roman';r.font.size=Pt(12);r.font.color.rgb=BLUE;r.bold=bold
def rh2(d,t):
    h=d.add_heading(t,level=2)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=RED
def bh3(d,t):
    h=d.add_heading(t,level=3)
    for r in h.runs:r.font.name='Times New Roman';r.font.color.rgb=BLUE
def sh(c,co):s=OxmlElement('w:shd');s.set(qn('w:fill'),co);s.set(qn('w:val'),'clear');c._tc.get_or_add_tcPr().append(s)
def sw(c,w):tcW=c._tc.get_or_add_tcPr();we=OxmlElement('w:tcW');we.set(qn('w:w'),str(int(w*567)));we.set(qn('w:type'),'dxa');tcW.append(we)
def tbl(d,headers,rows,widths):
    t=d.add_table(rows=1+len(rows),cols=len(headers));t.style='Table Grid';t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):sw(row.cells[ci],widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i];c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(10)
        sh(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci];c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:r.font.name='Times New Roman';r.font.size=Pt(10)

# ===== QT33: JAMA SCREEN MEDIA (3t+3h gốc) =====
d=md()
bl(d,'Tóm tắt bài QT-33',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Sử dụng phương tiện màn hình và sức khỏe tâm thần trẻ em và thanh thiếu niên: Phân tích thứ cấp từ thử nghiệm lâm sàng ngẫu nhiên \u00bb (Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial), do Schmidt-Persson J và cộng sự (2024), xuất bản trên JAMA Network Open (Q1, IF \u2248 13,0). Đan Mạch. 12 trang, 3 bảng, 3 hình.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Phân tích thứ cấp từ thử nghiệm lâm sàng ngẫu nhiên (RCT) \u2014 bằng chứng nhân quả MẠNH HƠN nghiên cứu quan sát. Nói cách khác, thay vì chỉ hỏi "VTN dùng màn hình nhiều có lo âu hơn không?", nghiên cứu này CAN THIỆP giảm screen time và đo tác động lên SKTT.',bold=True)
bl(d,'Tổng quan tài liệu cho thấy đa số NC về screen time là cắt ngang hoặc dọc quan sát. JAMA (IF = 13) công bố bằng chứng từ RCT \u2014 cấp độ bằng chứng cao nhất trong y học.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 1 gốc). Đặc điểm mẫu và thiết kế:', bold=True)
tbl(d,['Đặc điểm','Giá trị'],
    [['Thiết kế','Phân tích thứ cấp RCT'],
     ['Quốc gia','Đan Mạch'],
     ['Đối tượng','Trẻ em và thanh thiếu niên'],
     ['Can thiệp','Giảm thời gian sử dụng phương tiện màn hình'],
     ['Kết cục','Triệu chứng sức khỏe tâm thần'],
     ['Tạp chí','JAMA Network Open (Q1, IF \u2248 13,0)']],
    widths=[4.0,8.5])
d.add_paragraph()

bl(d,'Bảng 2. Hội tụ bằng chứng về screen time/MXH từ nhiều thiết kế:', bold=True)
tbl(d,['Nguồn','Thiết kế','Phát hiện','Cấp bằng chứng'],
    [['JAMA 2024 (bài này)','RCT (phân tích thứ cấp)','Giảm screen time \u2192 cải thiện SKTT','CAO NHẤT'],
     ['BJCP 2025 (QT22)','Dọc (longitudinal)','Screen time DỰ BÁO lo âu sau 1 năm','Cao'],
     ['Nature 2025 (QT27)','Cắt ngang + chẩn đoán lâm sàng','VTN có RLTT dùng MXH nhiều hơn','Trung bình-Cao'],
     ['Norway 2025 (QT21)','Xu hướng 13 năm','MXH giải thích xu hướng tăng','Trung bình'],
     ['Chen 2023 (bài 07)','Cắt ngang, n=63.205','Game OR = 5,00','Trung bình'],
     ['Hoàng Trung Học VN','Cắt ngang, n=8.473','Điện tử Beta = 0,176','Trung bình']],
    widths=[3.0,3.5,4.0,2.5])
d.add_paragraph()

bl(d,'Bảng 3. Kim tự tháp bằng chứng \u2014 vị trí các NC:', bold=True)
tbl(d,['Cấp','Thiết kế','Bài trong Đề tài'],
    [['1 (cao nhất)','RCT','JAMA 2024 (bài này)'],
     ['2','Dọc (longitudinal)','BJCP 2025 (QT22)'],
     ['3','Cắt ngang + chẩn đoán','Nature 2025 (QT27)'],
     ['4','Cắt ngang lớn','Chen 2023, Xu 2021, HTH 2025'],
     ['5','Tổng quan/meta','AJP 2024, BMC 2025, GBD 2025']],
    widths=[2.0,4.0,6.5])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : SÁU nguồn bằng chứng từ nhiều thiết kế (RCT, dọc, cắt ngang, xu hướng) và nhiều nước (Đan Mạch, Anh, Na Uy, Trung Quốc, Việt Nam) đều chỉ cùng hướng: screen time/MXH ảnh hưởng tiêu cực đến SKTT VTN. JAMA RCT cung cấp bằng chứng nhân quả mạnh nhất.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*Bằng chứng nhân quả từ RCT.* Đây là điểm mạnh vượt trội \u2014 giảm screen time CẢI THIỆN SKTT, không chỉ tương quan.')
bl(d,'*Ứng dụng cho VN.* Kết hợp với Hoàng Trung Học 2025 (điện tử Beta = 0,176) và Chen 2023 (game OR = 5,00), cung cấp cơ sở mạnh cho can thiệp quản lý thời gian màn hình tại trường học VN.')

rb(d,'Kết luận')
bl(d,'Dữ liệu từ RCT trên JAMA (IF = 13), cho thấy giảm thời gian sử dụng màn hình cải thiện SKTT ở trẻ em/VTN \u2014 bằng chứng nhân quả mạnh nhất hiện có, hội tụ với 5 nguồn bằng chứng khác từ nhiều thiết kế và quốc gia, gợi ý rằng can thiệp quản lý screen time là chiến lược phòng ngừa có cơ sở khoa học vững chắc.',bold=True)

rh2(d,'Phản biện')
bl(d,'JAMA Q1 IF = 13. Bằng chứng RCT \u2014 cấp cao nhất. Tuy nhiên, phân tích THỨ CẤP (không phải RCT thiết kế cho SKTT). Chỉ Đan Mạch \u2014 bối cảnh Bắc Âu khác châu Á. Cần đọc chi tiết cỡ mẫu và loại screen time (TV vs MXH vs game).')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'RCT giảm screen time tại trường VN. Phân biệt loại screen time (giải trí vs học tập). So sánh Đan Mạch vs châu Á.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50\u2b50 Rất cao. JAMA Q1 IF=13, RCT \u2014 bằng chứng nhân quả.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT33_JAMA_Screen.docx');sys.stderr.write('QT33 OK\n')

# ===== QT34: KOREA MH TRENDS (4t+1h gốc) =====
d=md()
bl(d,'Tóm tắt bài QT-34',bold=True)
rb(d,'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl(d,'Công trình \u00ab Xu hướng quốc gia về sức khỏe tâm thần thanh thiếu niên theo mức thu nhập tại Hàn Quốc, trước và sau COVID-19, 2006\u20132022 \u00bb (National Trends in Adolescents Mental Health by Income Level in South Korea, Pre- and Post-COVID-19, 2006\u20132022), xuất bản trên Nature Scientific Reports, 2024 (Q1, IF \u2248 4,6). Dữ liệu KYRBS (Korea Youth Risk Behavior Web-based Survey) quốc gia 2006\u20132022. 14 trang, 4 bảng, 1 hình.')

rb(d,'Phương pháp nghiên cứu')
bl(d,'Phân tích xu hướng 16 năm (2006\u20132022) từ KYRBS \u2014 khảo sát hành vi nguy cơ thanh niên quốc gia Hàn Quốc, phân tầng theo mức thu nhập gia đình và giai đoạn COVID-19. Nói cách khác, đây là phân tích xu hướng dài nhất trong Đề tài (16 năm), cho phép đánh giá tác động bất bình đẳng kinh tế lên SKTT VTN trước, trong và sau COVID-19.',bold=True)
bl(d,'KYRBS được Trung tâm Kiểm soát Dịch bệnh Hàn Quốc và Bộ Giáo dục triển khai hàng năm từ 2005, lấy mẫu cụm nhiều bậc đại diện quốc gia cho HS THCS + THPT.',bold=True)

rb(d,'Kết quả nghiên cứu định lượng')
bl(d,'Bảng 1 (từ Table 1-2 gốc). Xu hướng SKTT VTN Hàn Quốc 2006\u20132022:', bold=True)
tbl(d,['Chỉ số','Xu hướng trước COVID','Xu hướng sau COVID','Bất bình đẳng thu nhập'],
    [['Stress nhận thức','Giảm','TĂNG','Thu nhập thấp: 62,8% vs cao: 40,1%'],
     ['Buồn bã','Giảm','TĂNG (28,2%)','Khoảng cách MỞ RỘNG'],
     ['Ý tưởng tự tử','Giảm','TĂNG (13,9%)','Thu nhập thấp tệ hơn'],
     ['Cố tự tử','Giảm','TĂNG','Thu nhập thấp tệ hơn']],
    widths=[3.0,3.0,3.0,4.0])
d.add_paragraph()

bl(d,'Bảng 2. So sánh xu hướng Hàn Quốc với các nước:', bold=True)
tbl(d,['Nước','Giai đoạn','Xu hướng','Đặc thù'],
    [['Hàn Quốc (bài này)','2006\u20132022 (16 năm)','Giảm trước COVID, TĂNG sau','Bất bình đẳng thu nhập mở rộng'],
     ['Hoa Kỳ (QT23)','2013\u20132021 (8 năm)','Lo âu tăng gấp đôi','Hệ thống y tế công'],
     ['Na Uy (QT21)','2011\u20132024 (13 năm)','Tăng liên tục','Trường học + MXH'],
     ['Ireland (QT32)','2012\u20132019 (7 năm)','Tăng, nữ nhanh hơn','Cắt ngang lặp lại'],
     ['Việt Nam (VN14)','2021 vs 2023','41,5% \u2192 25,4%','Phục hồi sau COVID'],
     ['Toàn cầu (QT30)','1990\u20132021 (31 năm)','AAPC 0,84%','GBD 204 nước']],
    widths=[3.0,3.0,3.0,4.0])
d.add_paragraph()

bl(d,'Bảng 3. Bất bình đẳng thu nhập và SKTT:', bold=True)
tbl(d,['Thu nhập','Stress 2022','Buồn bã','Tự tử','So với VN'],
    [['Cao nhất','40,1%','Thấp hơn','Thấp hơn',''],
     ['Thấp nhất','62,8%','Cao hơn','Cao hơn',''],
     ['Chênh lệch','22,7 điểm','Mở rộng sau COVID','','VN: chưa có dữ liệu theo thu nhập']],
    widths=[2.5,2.5,2.5,2.5,3.0])
d.add_paragraph()

bh3(d,'Đối chiếu liên bài : Hàn Quốc là NƯỚC DUY NHẤT trong Đề tài cho thấy SKTT VTN CẢI THIỆN trước COVID (2006\u20132019) rồi XẤU ĐI sau COVID \u2014 mô hình ngược với phương Tây (tăng liên tục). Đáng chú ý, bất bình đẳng thu nhập MỞ RỘNG sau COVID \u2014 gợi ý COVID tác động nặng hơn lên nhóm nghèo. Tại VN, chưa có dữ liệu phân tầng theo thu nhập cho SKTT VTN \u2014 đây là GAP quan trọng.')

rb(d,'Nhận xét, phát hiện qua kết quả nghiên cứu')
bl(d,'*Xu hướng đảo chiều.* SKTT cải thiện 13 năm (2006\u20132019) rồi xấu đi sau COVID \u2014 mô hình độc đáo so với phương Tây (tăng liên tục). Có thể do chính sách SKTT học đường Hàn Quốc hiệu quả trước COVID.')
bl(d,'*Bất bình đẳng thu nhập.* Chênh lệch stress 22,7 điểm giữa nhóm giàu-nghèo là rất lớn. COVID làm khoảng cách MỞ RỘNG \u2014 nhóm nghèo bị ảnh hưởng nặng hơn.')
bl(d,'*Tự tử tăng sau COVID.* Hàn Quốc có tỷ lệ tự tử VTN cao nhất OECD \u2014 dữ liệu này xác nhận xu hướng đáng lo ngại.')

rb(d,'Kết luận')
bl(d,'Dữ liệu KYRBS quốc gia 16 năm, cho thấy SKTT VTN Hàn Quốc cải thiện trước COVID nhưng xấu đi đáng kể sau COVID, với bất bình đẳng thu nhập mở rộng (stress: 62,8% nhóm nghèo vs 40,1% nhóm giàu), gợi ý rằng (1) can thiệp SKTT có thể hiệu quả (giai đoạn cải thiện 2006\u20132019), nhưng (2) COVID đảo ngược tiến bộ, đặc biệt ở nhóm nghèo. Tại VN, cần phân tích SKTT theo thu nhập để đánh giá bất bình đẳng tương tự.',bold=True)

rh2(d,'Phản biện')
bl(d,'Nature Sci Rep Q1 IF = 4,6. KYRBS quốc gia 16 năm \u2014 xu hướng dài nhất trong Đề tài. Phân tầng theo thu nhập \u2014 ít NC nào làm. Tuy nhiên, chỉ Hàn Quốc (khác VN về hệ thống giáo dục, kinh tế). Đo stress/buồn bã (không tách lo âu riêng bằng GAD-7). KYRBS tự báo cáo trực tuyến \u2014 thiên lệch.')

rh2(d,'Hướng nghiên cứu tiếp theo')
bl(d,'Phân tích SKTT VTN VN theo thu nhập gia đình. So sánh xu hướng Hàn Quốc vs VN trước/sau COVID. Đánh giá chính sách SKTT học đường Hàn Quốc (giai đoạn cải thiện) để áp dụng cho VN.')

p=d.add_paragraph();r=p.add_run('Đánh giá: \u2b50\u2b50\u2b50\u2b50 Cao. Nature Q1, KYRBS 16 năm, bất bình đẳng thu nhập, mô hình đảo chiều COVID.');r.bold=True;r.font.name='Times New Roman';r.font.size=Pt(12)
d.save('QT34_Korea_Trends.docx');sys.stderr.write('QT34 OK\n')
