# -*- coding: utf-8 -*-
"""
Bài 03 v7 -> v8: sửa theo phản biện 2026.05.22 (Trường hợp 1 - chấp nhận đăng,
sửa nhẹ). Mọi chỗ sửa được TÔ ĐỎ để phản biện đối chiếu.

8 nhóm sửa:
1. Đồng bộ thuật ngữ "rối loạn lo âu" (bỏ "tổng quát") xuyên suốt VN+EN, trừ
   tên gốc thang đo DSM-5.
2. Tiêu đề EN.
3. Tóm tắt EN + Từ khóa EN.
4. Trích dẫn APA (2013) ở Đặt vấn đề.
5. Cách giới thiệu thang đo (Phương pháp).
6. Khai báo AI (đổi từ "cam kết" sang văn phong khai báo).
7. Bỏ ngôn ngữ nhân quả ở Bảng 6.
8. Xóa "Tuyên bố xung đột lợi ích" khỏi thân bài.
9. Bổ sung 3 TLTK Việt Nam đã xác minh.
"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn

IN = "bai-bao-khgdvn/bài-03/Công Thị Hằng_v7.docx"
OUT = "bai-bao-khgdvn/bài-03/Công Thị Hằng_v8.docx"

RED = RGBColor(0xFF, 0x00, 0x00)


def apply_replacements(text, replacements):
    """Áp lần lượt các thay thế, trả về list segments [(text, is_red)]."""
    segments = [(text, False)]
    for old, new in replacements:
        new_segments = []
        applied = False
        for seg_text, is_red in segments:
            if not applied and not is_red and old in seg_text:
                idx = seg_text.find(old)
                if idx > 0:
                    new_segments.append((seg_text[:idx], False))
                new_segments.append((new, True))
                if idx + len(old) < len(seg_text):
                    new_segments.append((seg_text[idx + len(old):], False))
                applied = True
            else:
                new_segments.append((seg_text, is_red))
        if not applied:
            print(f"  !! không tìm thấy: '{old[:40]}'")
        segments = new_segments
    # gộp các segment đỏ liền kề (nếu có) — không cần thiết, bỏ qua
    return segments


def set_segments(para, segments, size_pt=13):
    """Xóa hết run cũ, ghi mới theo segments."""
    for r in list(para._p.findall(qn("w:r"))):
        para._p.remove(r)
    for seg_text, is_red in segments:
        if not seg_text:
            continue
        run = para.add_run(seg_text)
        run.font.name = "Times New Roman"
        run.font.size = Pt(size_pt)
        if is_red:
            run.font.color.rgb = RED


def replace_para(para, new_text, size_pt=13, all_red=True):
    """Thay thế toàn bộ nội dung đoạn — toàn bộ tô đỏ."""
    set_segments(para, [(new_text, all_red)], size_pt)


d = Document(IN)
P = d.paragraphs

# ===== 1. Tiêu đề EN [0] =====
set_segments(P[0], [
    ("FAMILY ENVIRONMENT AND ", False),
    ("ANXIETY DISORDER", True),
    (" AMONG JUNIOR SECONDARY SCHOOL STUDENTS", False),
])
print("[0] Tiêu đề EN — đỏ ANXIETY DISORDER (bỏ GENERALIZED)")

# ===== 2. Tóm tắt EN [1] =====
abs_en = P[1].text
segs = apply_replacements(abs_en, [
    ("Generalized Anxiety Disorder (GAD)", "anxiety disorder"),
    ("predictor of GAD", "predictor of anxiety symptoms"),
    ("intervention of GAD", "intervention of anxiety disorder"),
])
set_segments(P[1], segs, size_pt=12)
print("[1] Tóm tắt EN — 3 chỗ GAD -> anxiety disorder/symptoms")

# ===== 3. Từ khóa EN [2] =====
set_segments(P[2], apply_replacements(P[2].text, [
    ("generalized anxiety disorder", "anxiety disorder"),
]), size_pt=12)
print("[2] Từ khóa EN — generalized anxiety disorder -> anxiety disorder")

# ===== 4. Tóm tắt VN [8] =====
set_segments(P[8], apply_replacements(P[8].text, [
    ("rối loạn lo âu tổng quát", "rối loạn lo âu"),
]), size_pt=12)
print("[8] Tóm tắt VN — bỏ 'tổng quát'")

# ===== 5. Từ khóa VN [9] =====
set_segments(P[9], apply_replacements(P[9].text, [
    ("rối loạn lo âu tổng quát", "rối loạn lo âu"),
]), size_pt=12)
print("[9] Từ khóa VN — bỏ 'tổng quát'")

# ===== 6. Trích dẫn APA [24] =====
set_segments(P[24], apply_replacements(P[24].text, [
    ("(Hiệp hội Tâm thần học Hoa Kỳ [American Psychiatric Association], 2013)",
     "(American Psychiatric Association — APA, 2013)"),
]))
print("[24] Trích dẫn APA — chuẩn hóa định dạng")

# ===== 7. Thang đo (Phương pháp) [38] =====
new_38 = ("(1) Thang đo mức độ nặng rối loạn lo âu (Severity Measure for "
          "Generalized Anxiety Disorder — DSM-5, dành cho trẻ 11–17 tuổi; "
          "APA, 2013), được sử dụng dưới dạng 5 mục phản ánh các biểu hiện "
          "cốt lõi của lo âu như lo lắng kéo dài, khó kiểm soát suy nghĩ lo âu, "
          "dự đoán tiêu cực và ảnh hưởng đến giấc ngủ. Thang đo sử dụng Likert "
          "4 mức (1 = không bao giờ; 4 = luôn luôn). Độ tin cậy được kiểm định "
          "bằng Cronbach's α và McDonald's ω; phân tích nhân tố khẳng định "
          "(CFA) được dùng để xác nhận giá trị cấu trúc của thang trong mẫu "
          "nghiên cứu.")
replace_para(P[38], new_38, size_pt=13, all_red=True)
print("[38] Thang đo (Phương pháp) — viết lại + nêu tên gốc DSM-5")

# ===== 8. Khai báo AI [47] =====
new_47 = ("Khai báo sử dụng công cụ trí tuệ nhân tạo. Trong quá trình hoàn "
          "thiện bản thảo, tác giả đã sử dụng công cụ trí tuệ nhân tạo dựa "
          "trên mô hình ngôn ngữ lớn (large language model) để rà soát hành "
          "văn, kiểm tra định dạng trích dẫn theo APA phiên bản thứ bảy và hỗ "
          "trợ sắp xếp danh mục tài liệu tham khảo. Thiết kế nghiên cứu, thu "
          "thập và phân tích dữ liệu, cũng như diễn giải kết quả là phần việc "
          "của tác giả. Số liệu trong bản thảo được đối chiếu với kết quả "
          "phân tích gốc. Công cụ trí tuệ nhân tạo không được ghi nhận là tác "
          "giả của bài viết.")
replace_para(P[47], new_47, size_pt=13, all_red=True)
print("[47] Khai báo AI — đổi từ 'cam kết' sang văn phong khai báo")

# ===== 9. Bảng 6 — bỏ ngôn ngữ nhân quả [59] =====
set_segments(P[59], apply_replacements(P[59].text, [
    ("cấu trúc lo âu tổng quát", "cấu trúc lo âu"),
    ("cho thấy đây là hệ quả thứ cấp của quá trình lo âu kéo dài, thay vì là "
     "biểu hiện cốt lõi.",
     "gợi ý đây có thể là biểu hiện ngoại vi hơn là biểu hiện cốt lõi trong "
     "cấu trúc lo âu của mẫu nghiên cứu."),
]))
print("[59] Bảng 6 — bỏ 'hệ quả thứ cấp' (ngôn ngữ nhân quả) + đồng bộ thuật ngữ")

# ===== 10-12. Bảng 8, 9, Ghi chú =====
for idx in [63, 64, 65]:
    set_segments(P[idx], apply_replacements(P[idx].text, [
        ("rối loạn lo âu tổng quát", "rối loạn lo âu"),
    ]))
    print(f"[{idx}] đồng bộ thuật ngữ")

# ===== 13-15. Thân bài [68], [69], [74] =====
for idx in [68, 69, 74]:
    set_segments(P[idx], apply_replacements(P[idx].text, [
        ("rối loạn lo âu tổng quát", "rối loạn lo âu"),
        ("lo âu tổng quát", "lo âu"),
    ]))
    print(f"[{idx}] đồng bộ thuật ngữ")

# ===== 16. Tuyên bố xung đột lợi ích [75] — XÓA, để lại note đỏ =====
new_75 = ("[ĐÃ XÓA THEO GÓP Ý PHẢN BIỆN: Tuyên bố xung đột lợi ích sẽ được "
          "khai báo qua biểu mẫu nộp bài riêng của Tạp chí, không in trong "
          "thân bài.]")
replace_para(P[75], new_75, size_pt=13, all_red=True)
print("[75] COI — XÓA + note đỏ")

# ===== 17. Bổ sung 3 TLTK Việt Nam =====
# Tìm vị trí chèn theo nội dung (tên cụ thể của đoạn ref TRƯỚC nó)
TLTK_VN = [
    # (chèn TRƯỚC ref có chứa đoạn này, nội dung mới)
    ("Hirsch, C. R.",
     "Hoàng Trung Học, & Nguyễn Thùy Dung. (2025). Levels of stress, anxiety, "
     "and depression in adolescents during and after the COVID-19 pandemic in "
     "Vietnam: A cross-sectional study. American Journal of Psychiatric "
     "Rehabilitation, 28(1), 360–367. https://doi.org/10.69980/ajpr.v28i1.105"),
    ("Hirsch, C. R.",
     "Phạm Thị Thu Hoa, Đỗ Thị Trang, Nguyễn Thị Liên, & Ngô Anh Vinh. "
     "(2024). Anxiety symptoms and coping strategies among high school "
     "students in Vietnam after COVID-19 pandemic: A mixed-method evaluation. "
     "Frontiers in Public Health, 12, Article 1232856. "
     "https://doi.org/10.3389/fpubh.2024.1232856"),
    ("Woodward, L. J.",
     "Đinh Thị Hồng Vân, Đỗ Thị Lê Hằng, & Phan Thị Mai Hương. (2021). "
     "Các yếu tố trường học ảnh hưởng đến lo âu của học sinh trung học cơ sở "
     "Việt Nam. Psychology and Education, 58(1), 883–894."),
]
for anchor_start, new_ref in TLTK_VN:
    target = None
    for p in d.paragraphs:
        if p.text.startswith(anchor_start):
            target = p
            break
    if target is None:
        print(f"  !! không tìm thấy anchor '{anchor_start}'")
        continue
    np = target.insert_paragraph_before()
    r = np.add_run(new_ref)
    r.font.name = "Times New Roman"
    r.font.size = Pt(13)
    r.font.color.rgb = RED
    print(f"  +TLTK đỏ chèn trước '{anchor_start[:30]}'")

# metadata
cp = d.core_properties
cp.author = ""
cp.last_modified_by = ""
d.save(OUT)
print(f"\nĐã lưu: {OUT}")
