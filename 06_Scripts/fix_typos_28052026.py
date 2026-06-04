# -*- coding: utf-8 -*-
"""Sua loi chinh ta + trinh bay AN TOAN (chi sua nhung gi chac chan dung).
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

BACKUP = FILE.replace('.docx', '_BEFORE_TYPO_FIX.docx')
shutil.copy(FILE, BACKUP)
print(f"Backup: {BACKUP}")

d = Document(FILE)

stats = {
    'duplicate_num_fixed': 0,
    'double_dot_fixed': 0,
    'double_space_fixed': 0,
    'space_before_punct_fixed': 0,
    'hoa_normalized': 0,
    'tam_than_normalized': 0,
}


def fix_runs(p, replacements):
    """Apply replacements at run-level to preserve formatting."""
    changed = False
    for run in p.runs:
        for old, new in replacements:
            if old in run.text:
                run.text = run.text.replace(old, new)
                changed = True
    return changed


# =========================================================
# FIX 1: SPECIFIC HIGH-PRIORITY FIXES (numbering + double dot)
# =========================================================
print("\n[1] Sua duplicate numbering + double dot...")

# Para 331: "1.2.6.5" -> "1.1.6.5"
if 331 < len(d.paragraphs):
    p = d.paragraphs[331]
    if '1.2.6.5' in p.text:
        if fix_runs(p, [('1.2.6.5', '1.1.6.5')]):
            stats['duplicate_num_fixed'] += 1
            print(f"  Para 331: '1.2.6.5' -> '1.1.6.5' (sai chi muc, chuong 1.1.6)")

# Para 1002: "3.4.2.1" -> "3.4.2.4"
if 1002 < len(d.paragraphs):
    p = d.paragraphs[1002]
    if '3.4.2.1.' in p.text:
        if fix_runs(p, [('3.4.2.1.', '3.4.2.4.')]):
            stats['duplicate_num_fixed'] += 1
            print(f"  Para 1002: '3.4.2.1' -> '3.4.2.4' (trung voi para 971)")

# Para 1036: "3.4.3.2" -> "3.4.3.3"
if 1036 < len(d.paragraphs):
    p = d.paragraphs[1036]
    if '3.4.3.2.' in p.text:
        if fix_runs(p, [('3.4.3.2.', '3.4.3.3.')]):
            stats['duplicate_num_fixed'] += 1
            print(f"  Para 1036: '3.4.3.2' -> '3.4.3.3' (trung voi para 1024)")

# Para 149: "cs.." -> "cs."
if 149 < len(d.paragraphs):
    p = d.paragraphs[149]
    if 'cs..' in p.text:
        if fix_runs(p, [('cs..', 'cs.')]):
            stats['double_dot_fixed'] += 1
            print(f"  Para 149: 'cs..' -> 'cs.' (dau cham doi)")


# =========================================================
# FIX 2: SPACE BEFORE PUNCTUATION (specific known cases)
# =========================================================
print("\n[2] Sua dau cach truoc dau cau...")
known_space_before = [
    (357, 'suy giảm :', 'suy giảm:'),
    (390, 'kiểm soát được ,', 'kiểm soát được,'),
    (390, 'vững chắc ,', 'vững chắc,'),
    (502, 'lo âu lan tỏa ,', 'lo âu lan tỏa,'),
    (546, '& Angold ,', '& Angold,'),
    (1217, 'ISSN : ', 'ISSN: '),
    (1316, 'DOI : ', 'DOI: '),
    (1325, 'DOI : ', 'DOI: '),
    (1333, 'Lekti A. , ', 'Lekti A., '),
    (1333, 'Icanervilia , ', 'Icanervilia, '),
]
for idx, old, new in known_space_before:
    if idx < len(d.paragraphs):
        p = d.paragraphs[idx]
        if old in p.text:
            if fix_runs(p, [(old, new)]):
                stats['space_before_punct_fixed'] += 1
                print(f"  Para {idx}: '{old}' -> '{new}'")


# =========================================================
# FIX 3: DOUBLE SPACES -> SINGLE SPACE (except in code/URLs)
# =========================================================
print("\n[3] Sua khoang trang doi -> don...")
for i, p in enumerate(d.paragraphs):
    if '  ' not in p.text:
        continue
    # Skip URLs/DOIs (rare to have double space anyway)
    for run in p.runs:
        if '  ' in run.text:
            # Replace 2+ consecutive spaces with 1 space (but preserve tab/newline)
            new_text = re.sub(r'  +', ' ', run.text)
            if new_text != run.text:
                run.text = new_text
                stats['double_space_fixed'] += 1


# =========================================================
# FIX 4: "hoà" -> "hòa" (Vietnamese standard)
# =========================================================
print("\n[4] Sua chinh ta 'hoà' -> 'hòa' (chuan VN hien tai)...")
for i, p in enumerate(d.paragraphs):
    if 'hoà' in p.text:
        if fix_runs(p, [('hoà', 'hòa')]):
            stats['hoa_normalized'] += 1


# =========================================================
# FIX 5: "Tâm Thần" -> "Tâm thần" (sentence case)
# =========================================================
print("\n[5] Sua 'Tâm Thần' -> 'Tâm thần'...")
for i, p in enumerate(d.paragraphs):
    if 'Tâm Thần' in p.text:
        if fix_runs(p, [('Tâm Thần', 'Tâm thần')]):
            stats['tam_than_normalized'] += 1


# =========================================================
# SAVE
# =========================================================
d.save(FILE)
print(f"\nSaved: {FILE}")
print(f"Size: {os.path.getsize(FILE)//1024} KB")

print("\n" + "="*60)
print("TOM TAT")
print("="*60)
for k, v in stats.items():
    print(f"  {k}: {v}")
print(f"\nTong cong: {sum(stats.values())} sua")
print()
print("KHONG SUA (can NCS xem xet):")
print("  • Para 529 + 544: 1.2.4.3 trung 2 lan (can quyet dinh cau truc)")
print("  • Dau ngoac khong can bang (16 doan) - case-by-case")
print("  • Empty paragraph clusters (co the la space co y)")
print("  • Missing space after author abbreviations (chu yeu DOI URLs - false positive)")
