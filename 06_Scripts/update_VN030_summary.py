# -*- coding: utf-8 -*-
"""Cap nhat VN030 summary de match voi full translation."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name='Times New Roman'; s.font.size=Pt(12)
    s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.5
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.5)
        sec.left_margin=Cm(3); sec.right_margin=Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='Times New Roman'; r.font.color.rgb=RGBColor(0,0,0)
def P(doc, text, bold=False, italic=False, size=12):
    p = doc.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    t.autofit=False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)

# ============================================================
doc = make_doc()
H(doc, 'Tóm tắt bài VN030 — Happy House: RCT cluster can thiệp SKTT trường học cho VTN Hà Nội', level=1)
P(doc, 'Cập nhật từ full text PMC10643236 (12/04/2026).', italic=True)

P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "School-based universal mental health promotion intervention for adolescents '
       'in Vietnam: Two-arm, parallel, controlled trial" của Tran, T.D., Nguyen, H., Shochet, I., '
       'Nguyen, N., La, N., Wurfl, A., Orr, J., Nguyen, H., Stocker, R. & Fisher, J. (2023), '
       'đăng trên Cambridge Prisms: Global Mental Health vol. 10, e58 (10/2023). Link PMC: '
       'https://pmc.ncbi.nlm.nih.gov/articles/PMC10643236/. Đơn vị hợp tác Việt – Úc: Monash '
       'University (Melbourne), Hanoi University of Public Health, Queensland University of '
       'Technology (Brisbane). Khách thể: 1.084 học sinh lớp 10 (531 can thiệp + 552 đối chứng) '
       'tại 8 trường THPT Hà Nội (4+4). Tỷ lệ tham gia 96,1 % (1.084 từ 1.128 đủ điều kiện). '
       'Tuổi 15-16. Trial registration: ACTRN12620000088943. Funding: NHMRC Úc (GNT1158429) + '
       'NAFOSTED Việt Nam (NHMRC.108.01-2018.02).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Thiết kế cluster controlled (two-arm parallel) — phân bổ cấp TRƯỜNG (cluster non-randomized '
       'do yêu cầu hành chính). 8 trường chọn dựa trên sẵn sàng tham gia + cân bằng đặc điểm '
       'kinh tế xã hội + địa lý (đô thị vs ngoại ô). 4 trường can thiệp + 4 trường đối chứng.')
P(doc, 'Can thiệp Happy House: phiên bản thích ứng văn hoá Việt Nam của chương trình RAP-A '
       '(Resourceful Adolescent Program), 6 buổi mỗi buổi 90 phút trong 6 tuần liền. Nội dung: '
       'CBT + IPT thông qua 6 chủ đề (sức mạnh bản thân, suy nghĩ tích cực, quản lý cảm xúc, '
       'giải quyết vấn đề, mạng lưới hỗ trợ, ôn tập + cam kết). Thay thế môn Giáo dục Công dân.')
P(doc, 'Người cung cấp: Giáo viên Giáo dục Công dân được đào tạo 3 ngày bởi Monash + HUPH. '
       'Mỗi buổi có giáo viên chính + 1-2 trợ giảng để hỗ trợ lớp đông (40-45 HS). 95 % HS '
       'hoàn thành toàn bộ 6 buổi, 100 % fidelity tự đánh giá.')
P(doc, 'Đo lường: CESD-R (Center for Epidemiologic Studies Depression Revised, 20 mục, thang '
       '0-80, ngưỡng ≥ 16) là kết cục chính. Kết cục phụ: MHC-SF (phúc lợi), CSES (tự hiệu '
       'quả ứng phó), school connectedness. Thời điểm: ban đầu (10/2020), 2 tuần sau (12/2020), '
       '6 tháng sau (05/2021, online do COVID-19). Phân tích: mixed-effects models với cụm '
       'lớp ngẫu nhiên, ITT, kiểm soát baseline.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'CESD-R ≥ 16 (kết cục chính):', bold=True)
table(doc,
    ['Thời điểm', 'Đối chứng (%)', 'Can thiệp (%)', 'aOR (KTC 95 %)', 'p', "Cohen's d"],
    [
        ['2 tuần sau', '28,6 % (155/542)', '23,9 % (125/523)', '0,56 (0,36–0,88)', '0,011', '0,11'],
        ['6 tháng sau', '29,2 % (157/537)', '26,3 % (138/524)', '0,75 (0,51–1,09)', '0,132', '—'],
    ],
    widths=[2.5, 3.3, 3.3, 3.0, 1.4, 2.0])

P(doc, 'Kết cục phụ:', bold=True)
P(doc, '• Phúc lợi tâm lý (MHC-SF): d = 0,13 (p < 0,05) ở 2 tuần, không còn ý nghĩa ở 6 tháng')
P(doc, '• CSES ứng phó vấn đề: d = 0,17–0,26 DUY TRÌ ở cả 2 tuần và 6 tháng *')
P(doc, '• CSES tìm kiếm hỗ trợ XH: DUY TRÌ ở 6 tháng *')
P(doc, '• CSES điều tiết cảm xúc: không còn ý nghĩa sau Bonferroni')
P(doc, '• School connectedness: không có tác động')
P(doc, '(* = có ý nghĩa thống kê sau hiệu chỉnh Bonferroni)')

P(doc, 'Các con số quan trọng:', bold=True)
P(doc, '• 1.084 HS lớp 10 (96,1 % từ 1.128 đủ điều kiện)')
P(doc, '• 531 can thiệp + 552 đối chứng')
P(doc, '• Baseline CESD-R: 11,4 ± 12,2 (ĐC) vs 12,0 ± 12,0 (CT), p = 0,729')
P(doc, '• Tỷ lệ trầm cảm lâm sàng ban đầu: 25,5 % (cả 2 nhóm)')
P(doc, '• Mất mẫu: < 3 % ở mọi thời điểm')

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Cluster controlled trial ĐẦU TIÊN tại Việt Nam về can thiệp SKTT trường '
       'học cho VTN. Cỡ mẫu rất lớn (n = 1.084), tỷ lệ tham gia 96,1 %, mất mẫu < 3 %. RAP-A '
       'có cơ sở bằng chứng quốc tế mạnh từ Úc, New Zealand, Mauritius. Mô hình hợp tác Việt '
       '– Úc (Monash–QUT–HUPH) uy tín. Giáo viên cung cấp — phù hợp LMIC. Đăng ký thử nghiệm '
       'trước + tài trợ quốc tế.', bold=True)

P(doc, 'Hạn chế: (1) Can thiệp UNIVERSAL (không phải TARGETED) → hiệu ứng nhỏ d = 0,11; '
       '(2) FADE-OUT ở 6 tháng — gợi ý cần BOOSTER SESSIONS; (3) Đo TRẦM CẢM chính, KHÔNG '
       'đo LO ÂU riêng; (4) Chỉ lớp 10 Hà Nội — chưa đại diện VN; (5) Phân bổ NHÓM không '
       'ngẫu nhiên (có thể bias giữa trường can thiệp vs đối chứng); (6) KHÔNG có active '
       'control; (7) COVID-19 confounding thời gian theo dõi.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, '(1) Chuyển sang mô hình TARGETED (HS có GAD-7 ≥ 8 hoặc CESD-R ≥ 16) để giảm dilution '
       'và có hiệu ứng lớn hơn. (2) Thêm BOOSTER SESSIONS ở 3 và 6 tháng để duy trì hiệu quả. '
       '(3) Đo CẢ TRẦM CẢM + LO ÂU + coping self-efficacy. (4) Mở rộng địa bàn (TPHCM, Đà '
       'Nẵng, miền Trung nông thôn, DTTS). (5) RANDOMIZE cấp trường để có RCT thuần. (6) So '
       'sánh giáo viên vs chuyên gia tâm lý vs y tá học đường cung cấp.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. Cluster controlled trial ĐẦU TIÊN tại VN cho VTN, '
       'mẫu lớn 1.084, tỷ lệ tham gia 96,1 %, tác giả Monash–QUT–HUPH. Hạn chế: universal '
       'intervention với hiệu ứng nhỏ (d = 0,11), không đo lo âu riêng, không randomized. '
       'PHẢI trích dẫn trong mọi báo cáo can thiệp VN từ nay.', bold=True)

# Save new version
out_path = os.path.join(TT_DIR, 'VN030_Tran_HappyHouse_Cambridge_2023.docx')
doc.save(out_path)
d = Document(out_path)
chars = sum(len(p.text) for p in d.paragraphs)
print(f'Saved: {os.path.basename(out_path)}')
print(f'Chars: {chars}, Tables: {len(d.tables)}')

# Remove old short version
old_path = os.path.join(TT_DIR, 'VN030_Tran_HappyHouse_RCT_GMH_2023.docx')
if os.path.exists(old_path):
    os.remove(old_path)
    print(f'Removed old: VN030_Tran_HappyHouse_RCT_GMH_2023.docx')
