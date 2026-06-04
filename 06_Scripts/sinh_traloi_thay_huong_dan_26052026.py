# -*- coding: utf-8 -*-
"""Sinh file 04 - Tra loi thay huong dan ngay 26/05/2026.
Tra loi 4 y cua thay:
  1. Chien luoc "de hoi dai de HD rut"
  2. Vai tro Bai 1 + Bai 2 trong giai trinh
  3. Bien phap ung pho co de vao YTBV khong?
  4. Ke hoach Tong quan ngan upper part
File: Luan an TS/04_TraLoi_ThayHuongDan_v1_26052026.docx"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '04_TraLoi_ThayHuongDan_v1_26052026.docx')

# ============================================================
# HELPERS
# ============================================================
def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.0); sec.right_margin = Cm(2.0)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)

def P(doc, text, indent=True, italic=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    r.italic = italic; r.bold = bold
    return p

def P_quote(doc, text):
    """Doan trich loi thay - in nghieng, le trai, khong indent."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.right_indent = Cm(1.0)
    r = p.add_run('"' + text + '"')
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.italic = True

# ============================================================
# BUILD DOC
# ============================================================
doc = doc_init()

# Title
H(doc, 'TRẢ LỜI THẦY HƯỚNG DẪN', 1)
P(doc, '(Sau khi nhận tin nhắn của thầy về chiến lược bản nộp và Bài 1 + Bài 2 — 26/05/2026)', indent=False, italic=True)
doc.add_paragraph()

# Lời mở
P(doc, 'Kính gửi thầy,')
P(doc, 'Em xin cảm ơn thầy đã có những phản hồi rất kịp thời và định hướng cụ thể cho bản nộp. Em xin được trả lời lần lượt bốn ý của thầy như sau.')
doc.add_paragraph()

# ============================================================
# 1. CHIẾN LƯỢC "ĐỂ DÀI"
# ============================================================
H(doc, '1. Về chiến lược "để hơi dài để Hội đồng yêu cầu rút ngắn"', 2)

P(doc, 'Em đồng tình với chiến lược của thầy. Việc giữ phần tổng quan và cơ sở lý luận ở mức độ rộng vừa phải có hai lợi ích thực tế. Thứ nhất, Hội đồng có vật liệu rõ ràng để chỉ ra "phần nào cần rút" — đây là dạng yêu cầu dễ thực hiện và không ảnh hưởng tới khung lý thuyết hay phương pháp. Thứ hai, nếu nội dung quá mỏng, Hội đồng có thể yêu cầu bổ sung thêm khảo sát thực nghiệm hoặc phân tích phụ — những việc cần nhiều thời gian hơn nhiều so với cắt gọn.')

P(doc, 'Em xin được bổ sung hai lưu ý nhỏ để chiến lược này không bị "phản tác dụng". Thứ nhất, các phần liên quan đến phương pháp (Chương 2) và thống kê (Chương 3, mục 3.4) nên giữ ở mức gọn và rõ — vì đây là chỗ phản biện đã chỉ ra nhiều điểm cần làm rõ, nếu để dài có thể bị soi sâu hơn. Thứ hai, phần Kết luận và Khuyến nghị nên trình bày dứt khoát, không trộn lẫn bàn luận — phản biện đã chỉ ra điều này ở mục 48. Như vậy "để dài" chủ yếu nên tập trung ở tổng quan và cơ sở lý luận, là phần dễ điều chỉnh khi Hội đồng góp ý.')
doc.add_paragraph()

# ============================================================
# 2. VAI TRÒ BÀI 1 + BÀI 2
# ============================================================
H(doc, '2. Vai trò của Bài 1 và Bài 2 trong giải trình với Hội đồng', 2)

P(doc, 'Hai bài viết em vừa gửi thầy đã được thiết kế làm hai chương ngắn độc lập, đồng thời cũng là vật liệu hỗ trợ giải trình ở những phần dễ bị Hội đồng chất vấn. Cụ thể:')

P(doc, 'Bài 1 — Yếu tố nguy cơ rối loạn lo âu ở học sinh THCS (17.485 từ; 21 tài liệu tham khảo đã được đối chiếu DOI/PubMed): hệ thống hóa năm nhóm yếu tố nguy cơ (áp lực học tập, nghiện điện thoại và mạng xã hội, bắt nạt thể chất, bắt nạt bằng lời và bắt nạt mạng, lòng tự trọng thấp), kèm theo phân tích tương tác giữa các nhóm. Bài này trả lời trực tiếp các câu hỏi của Hội đồng (nếu có) về: (i) khung phân tích yếu tố nguy cơ — tiếp cận sinh thái xã hội ba tầng; (ii) năm khoảng trống nghiên cứu trong nước; (iii) lý do chọn các yếu tố nguy cơ cụ thể được khảo sát trong luận án.')

P(doc, 'Bài 2 — Khoảng trống can thiệp rối loạn lo âu ở học sinh THCS (13.839 từ; 21 tài liệu tham khảo): hệ thống hóa ba kênh can thiệp (lâm sàng, trường học, số) trong khung phân tầng ba tầng của Mrazek và Haggerty (1994), cùng mô hình stepped care của Bower và Gilbody (2005). Bài này đặc biệt hữu ích để giải trình phần 3.7 của luận án — em đã thiết kế bài này nhằm cung cấp cơ sở lý luận chi tiết cho khung chương trình tập huấn phòng ngừa, một điểm mà phản biện đã chất vấn về thiếu căn cứ.')

P(doc, 'Khi gửi cho Hội đồng, em xin đề xuất nộp cả hai bài kèm theo luận án ở dạng tài liệu tham chiếu. Trong luận án chính, các phần tổng quan về yếu tố nguy cơ và can thiệp sẽ trình bày tóm tắt, có dẫn rõ rằng phần lập luận đầy đủ nằm ở hai bài đính kèm. Cách này vừa giảm dung lượng tổng quan trong luận án (đúng theo gợi ý của thầy), vừa giúp Hội đồng có vật liệu chi tiết nếu họ muốn đào sâu một điểm cụ thể.')
doc.add_paragraph()

# ============================================================
# 3. CÂU HỎI VỀ BIỆN PHÁP ỨNG PHÓ
# ============================================================
H(doc, '3. Về câu hỏi của thầy: "Các biện pháp ứng phó có nên để vào yếu tố bảo vệ không?"', 2)

P_quote(doc, 'Các biện pháp ứng phó có nên để vào yếu tố bảo vệ không Em?')

P(doc, 'Em xin trả lời: nên đặt biện pháp ứng phó vào nhóm yếu tố bảo vệ, với điều kiện làm rõ bản chất "tùy chiến lược" của nó. Em xin trình bày ba lý do.')

P(doc, 'Thứ nhất, về lý thuyết kinh điển. Lazarus và Folkman (1984) trong khung lý thuyết về ứng phó phân biệt hai dạng chính: ứng phó tập trung vào vấn đề (problem-focused) và ứng phó tập trung vào cảm xúc (emotion-focused). Cả hai dạng đều có thể là bảo vệ khi được sử dụng adaptive (chủ động, hướng tới giải quyết hoặc điều tiết hiệu quả). Ngược lại, ứng phó dạng né tránh, kìm nén cảm xúc hoặc phủ nhận lại làm trầm trọng triệu chứng. Như vậy, "biện pháp ứng phó" như một biến tổng thường được xếp vào nhóm yếu tố bảo vệ trong tài liệu nghiên cứu chính thống.')

P(doc, 'Thứ hai, về bằng chứng định lượng quốc tế. Phân tích tổng hợp của Compas và cộng sự (2017) trên 80.850 trẻ em — vị thành niên qua 212 nghiên cứu, công bố trên Psychological Bulletin 143(9), xác nhận rằng ứng phó tập trung vào vấn đề và đánh giá lại nhận thức có liên hệ ngược chiều và bền vững với triệu chứng lo âu và trầm cảm. Trong khi đó, ứng phó né tránh và kìm nén cảm xúc có liên hệ thuận chiều. Phát hiện này đã được trích dẫn ở Bài 1 — Yếu tố nguy cơ (mục 5.2) khi nói về lòng tự trọng và ứng phó như những cấu trúc bảo vệ.')

P(doc, 'Thứ ba, về cách trả lời phản biện ở các mục 8, 14, 28 và 49. Thầy phản biện chỉ ra rằng "biện pháp đối phó" được trình bày không nhất quán trong luận án — có chỗ là yếu tố bảo vệ, có chỗ lại như biến phụ thuộc. Em xin được đề xuất hướng giải quyết như sau. Trong Cơ sở lý luận (Chương 1), đưa ứng phó vào cùng nhóm với lòng tự trọng, gắn bó trường học và hỗ trợ xã hội — bốn cấu trúc bảo vệ kinh điển. Trong Phương pháp (Chương 2), nêu rõ thang đo (ví dụ Brief COPE hoặc phiên bản rút gọn phù hợp) phân tách hai chiều ứng phó adaptive và maladaptive. Trong Kết quả (Chương 3), phân tích riêng từng chiều và bàn luận hệ số β tương ứng — chiều adaptive kỳ vọng β âm với điểm lo âu, chiều maladaptive kỳ vọng β dương. Cách trình bày này vừa giữ được khung lý thuyết "ứng phó là yếu tố bảo vệ" như thầy gợi ý, vừa trả lời được câu hỏi của phản biện về tính nhất quán biến.')

P(doc, 'Em cũng xin lưu ý điểm phản biện mục 29 nêu: "một yếu tố chỉ được xem là bảo vệ hay nguy cơ dựa trên số liệu đo đạc". Quan điểm này đúng ở phương pháp luận. Tuy nhiên, ở cấp độ trình bày tổng quan và cơ sở lý luận, việc phân nhóm theo các quy ước có sẵn trong tài liệu vẫn là cần thiết — miễn là luận án nêu rõ "nhóm phân loại lý thuyết" và "nhóm phân loại thực nghiệm" có thể không trùng nhau, và kết quả thực nghiệm sẽ là tiêu chí cuối cùng. Em xin viết một đoạn ngắn ở đầu mục Yếu tố ảnh hưởng (Chương 1) làm rõ điểm này.')
doc.add_paragraph()

# ============================================================
# 4. KẾ HOẠCH VIẾT TỔNG QUAN NGẮN
# ============================================================
H(doc, '4. Kế hoạch viết Tổng quan ngắn (phần trên của Chương 1)', 2)

P(doc, 'Theo gợi ý của thầy, em sẽ viết lại phần trên của tổng quan, đi thẳng vào các nội dung thầy đã liệt kê, để ghép với Bài 1 và Bài 2 đã hoàn thành. Mục tiêu dung lượng khoảng 8.000-10.000 từ — gọn hơn đáng kể so với 21.000 từ hiện tại nhưng vẫn đủ "dày" để Hội đồng tập trung yêu cầu rút gọn ở đây, tránh việc Hội đồng đi tìm điểm yếu khác bắt em làm lại. Đây là vận dụng chính chiến lược "để hơi dài để Hội đồng yêu cầu rút" mà thầy đã đưa ra ở ý đầu.')

P(doc, 'Cách trình bày theo CHỦ ĐỀ thay vì theo từng công trình — tương đồng với cấu trúc của Bài 1 và Bài 2 mà em đã viết. Thay vì lần lượt giới thiệu "Công trình của X (năm) về..." rồi tóm tắt từng nghiên cứu riêng lẻ, em sẽ gom các phát hiện cùng chủ đề vào một mạch lập luận, dẫn chứng nhiều nghiên cứu trong cùng đoạn để tổng hợp một xu hướng. Cách viết này có ba lợi điểm: (i) khắc phục đúng phản biện ở Mục 10 — "NCS trình bày liệt kê từng công trình, chưa tổng quát hóa được"; (ii) gọn dung lượng tự nhiên (không cần nhắc lại bối cảnh từng nghiên cứu); (iii) tạo sự nhất quán xuyên suốt với hai bài tổng quan đính kèm về Yếu tố nguy cơ và Can thiệp.')

P(doc, 'Cấu trúc đề xuất gồm bảy tiểu mục như sau. Tiểu mục 1 — Tỷ lệ phổ biến (khoảng 1.500 từ): tổng hợp các số liệu trọng tâm cả quốc tế và trong nước — Phạm Thị Thu Hoa và cộng sự (2024) trên 3.910 học sinh THPT Hà Nội (tỷ lệ 40,6%), V-NAMHS 2022 với 5.996 cặp cha mẹ — vị thành niên (rối loạn lo âu DSM-5 = 2,3%), Hoàng Trung Học và Nguyễn Thùy Dung (2025) với mẫu lớn nhất hiện có, Chen và cộng sự (2023) — 63.205 học sinh Trung Quốc với 13,9%; có phân tích khác biệt giới tính và công cụ đo (GAD-7, DASS-21, RCADS, SCARED). Tiểu mục 2 — Mức độ (khoảng 1.000 từ): phân tầng nhẹ/vừa/nặng theo các công cụ phổ biến, khác biệt giới và khối lớp. Tiểu mục 3 — Biểu hiện (khoảng 1.500 từ): ba dạng rối loạn lo âu lan tỏa, lo âu xã hội, lo âu chia ly theo phân loại DSM-5, kèm dẫn chứng định lượng cho từng dạng. Tiểu mục 4 — Yếu tố nguy cơ (khoảng 800 từ tóm tắt năm nhóm yếu tố nguy cơ chính) với dẫn rõ "phân tích đầy đủ xem Bài 1 đính kèm". Tiểu mục 5 — Yếu tố bảo vệ (khoảng 2.500 từ): đây là phần em sẽ viết chi tiết nhất trong tổng quan, vì chưa có bài riêng cho chủ đề này — bao gồm lòng tự trọng, gắn bó trường học, hỗ trợ của cha mẹ, hỗ trợ từ bạn bè, và biện pháp ứng phó (theo cách phân tích ở mục 3 trên). Tiểu mục 6 — Can thiệp (khoảng 800 từ tóm tắt ba kênh can thiệp) với dẫn "phân tích đầy đủ xem Bài 2 đính kèm". Tiểu mục 7 — Nhận xét chung và khoảng trống ở Việt Nam (khoảng 800 từ).')

P(doc, 'Cách trình bày này vẫn ngắn hơn rõ rệt so với 21.000 từ hiện tại (giảm hơn nửa), giữ được tỉ lệ hợp lý với phần thực trạng (Chương 3 đang khoảng 18.000 từ), đồng thời đủ "thân" để Hội đồng có chỗ để chỉ trỏ và yêu cầu rút thêm. Khi nào bản nộp đã qua phản biện kín và đến giai đoạn hoàn thiện, em sẽ mở rộng lại các phần Hội đồng yêu cầu giữ nguyên.')

P(doc, 'Em xin được làm việc với thầy về bản phác thảo của tổng quan ngắn trong vài ngày tới, trước khi tích hợp chính thức vào luận án.')
doc.add_paragraph()

# Closing
P(doc, 'Em xin cảm ơn thầy.', indent=False)
P(doc, 'Trân trọng,', indent=False)
P(doc, '', indent=False)
P(doc, 'Công Thị Hằng', indent=False, bold=True)

# ============================================================
# LUU
# ============================================================
doc.save(OUT)

from docx import Document as D
d = D(OUT)
w = sum(len(p.text.split()) for p in d.paragraphs)
n_p = sum(1 for p in d.paragraphs if p.text.strip())
print(f"Da luu: {OUT}")
print(f"Tu: {w} | Doan: {n_p} | Size: {os.path.getsize(OUT) // 1024}KB")
