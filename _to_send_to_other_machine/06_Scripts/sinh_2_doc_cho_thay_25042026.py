# -*- coding: utf-8 -*-
"""Sinh 2 file docx trả lời thầy (25/04/2026):
  1) McDonald's omega là gì?
  2) Bài 8 (Wen 2020) — đính chính chiều giới tính & diễn giải OR
Dựa trên truy vấn RAG (rag_db_full collection lo_au_full) + KG (kg_data) + kiến thức tâm trắc học chuẩn.
Một lần dùng.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao'
RED  = RGBColor(0xC0, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)

def new_doc():
    d = Document()
    s = d.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in d.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return d

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)

def tbl(doc, headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

def title(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(16); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'

def subtitle(doc, text):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'

def h(doc, text):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'

def nr(doc, text, bold=False, size=13, color=None, italic=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color
    return p

def mixed(doc, parts):
    """parts: list of (text, dict opts: bold, color, italic, size)."""
    p = doc.add_paragraph()
    for text, opts in parts:
        r = p.add_run(text); r.font.name = 'Times New Roman'
        r.font.size = Pt(opts.get('size', 13))
        r.bold = opts.get('bold', False); r.italic = opts.get('italic', False)
        if 'color' in opts: r.font.color.rgb = opts['color']
    return p

# =====================================================================
# DOC 1 — McDonald's omega
# =====================================================================
def make_doc_omega():
    d = new_doc()
    title(d, "HỆ SỐ McDONALD'S OMEGA (ω) LÀ GÌ?")
    subtitle(d, "Giải thích nhanh cho thầy — dành cho người làm lâm sàng/nghiên cứu tâm lý")
    nr(d, "Trợ lý nghiên cứu — 25/04/2026 — soạn từ truy vấn RAG (lo_au_full) + tâm trắc học chuẩn.",
       size=10, italic=True, color=GRAY)

    h(d, "1. Trả lời ngắn (1 đoạn)")
    nr(d,
       "McDonald's omega (ký hiệu ω) là một hệ số đánh giá độ tin cậy nội tại "
       "(internal consistency reliability) của một thang đo tâm lý — giống như "
       "Cronbach's alpha (α), nhưng dựa trên mô hình nhân tố (factor model) thay vì giả định "
       "tau-tương đương. Nói nôm na: ω cho biết các câu hỏi (item) trong thang đo có cùng "
       "đo một cấu trúc tiềm ẩn (ví dụ \"lo âu\") đến mức nào, sau khi mô hình hoá sự đóng "
       "góp khác nhau của từng item bằng hệ số tải nhân tố (factor loading) λᵢ.")

    h(d, "2. Công thức cơ bản (omega total — McDonald, 1999)")
    nr(d, "ω = (Σλᵢ)² / [ (Σλᵢ)² + Σ(1 − λᵢ²) ]", bold=True)
    nr(d, "Trong đó λᵢ là factor loading chuẩn hoá của item thứ i lên nhân tố chung. "
          "Tử số là phương sai \"thật\" do nhân tố chung giải thích; mẫu số là tổng phương sai "
          "(thật + sai số đặc thù của từng item). Một số phần mềm (R: psych::omega, JASP, jamovi) "
          "phân biệt thêm omega hierarchical (ωₕ) — chỉ tính phần do nhân tố chung sau khi loại "
          "các nhân tố nhóm.")

    h(d, "3. So sánh nhanh ω vs α")
    tbl(d,
        ['Tiêu chí', "Cronbach's alpha (α)", "McDonald's omega (ω)"],
        [
            ['Cơ sở lý thuyết', 'Lý thuyết trắc nghiệm cổ điển',
             'Mô hình nhân tố (CFA / EFA)'],
            ['Giả định mấu chốt', 'Tau-tương đương: mọi item đóng góp như nhau (λᵢ bằng nhau)',
             'Cho phép factor loading khác nhau giữa các item — sát thực tế hơn'],
            ['Khi giả định bị vi phạm', 'α thường ƯỚC LƯỢNG THẤP độ tin cậy thực',
             'ω vẫn cho ước lượng chính xác'],
            ['Cần chạy CFA trước?', 'Không',
             'Có (cần có hệ số tải nhân tố)'],
            ['Khuyến cáo hiện nay', 'Vẫn dùng nhưng hạn chế',
             'Nhiều tạp chí (JCCAP, Psychol Methods…) đã ưu tiên ω'],
        ],
        [3.5, 5.5, 5.5])
    nr(d, "(Tham khảo: McDonald 1999; Dunn, Baguley & Brunsden 2014; Hayes & Coutts 2020 — "
          "« Use Omega Rather than Cronbach's Alpha »).", size=10, italic=True, color=GRAY)

    h(d, "4. Ngưỡng diễn giải (giống α)")
    tbl(d,
        ['Khoảng giá trị ω', 'Diễn giải'],
        [
            ['ω ≥ 0,90', 'Xuất sắc — nhưng cảnh giác item trùng lặp nội dung (redundancy)'],
            ['0,80 ≤ ω < 0,90', 'Tốt — phù hợp nghiên cứu và sàng lọc lâm sàng'],
            ['0,70 ≤ ω < 0,80', 'Chấp nhận được — đủ dùng cho nghiên cứu'],
            ['0,60 ≤ ω < 0,70', 'Nghi ngờ — chỉ dùng thăm dò'],
            ['ω < 0,60', 'Không chấp nhận — cần xem lại cấu trúc thang đo'],
        ],
        [4.5, 10.0])

    h(d, "5. Khi nào nên ưu tiên ω hơn α?")
    nr(d, "• Hệ số tải nhân tố giữa các item KHÔNG đồng đều (rất phổ biến trong DASS-21, "
          "GAD-7, PHQ-9 khi áp dụng trên mẫu mới hoặc qua dịch ngôn ngữ).")
    nr(d, "• Thang đo có cấu trúc đa chiều (multidimensional) — ví dụ DASS-21 có 3 thang con "
          "Trầm cảm/Lo âu/Stress; nên báo ω cho từng thang con.")
    nr(d, "• Cỡ mẫu nhỏ hoặc phương sai item chênh lệch lớn.")
    nr(d, "• Khi cần báo cáo độ tin cậy của một thang đo đã có CFA xác nhận.")

    h(d, "6. Liên hệ với corpus 35+ bài lo âu trong dự án")
    mixed(d, [
        ("Truy vấn RAG (collection ", {}),
        ("lo_au_full", {'italic': True}),
        (", 47 chunks) cho 4 biến thể câu hỏi về omega đều ", {}),
        ("KHÔNG", {'bold': True, 'color': RED}),
        (" trả về chunk định nghĩa trực tiếp — vì các bài trong corpus tập trung vào "
         "tỉ lệ lo âu, yếu tố nguy cơ, can thiệp; phần tâm trắc học chỉ báo cáo giá trị α "
         "hoặc ω cho thang đo đã dùng (ví dụ ", {}),
        ("Hoa 2024 báo Cronbach α = 0,916 cho GAD-7 tại Hà Nội", {'italic': True}),
        ("). Vì vậy phần định nghĩa ở trên dựa trên tài liệu tâm trắc học chuẩn "
         "(McDonald 1999; Hayes & Coutts 2020), RAG đóng vai trò bổ trợ.", {}),
    ])

    h(d, "7. Tóm gọn để nói nhanh với sinh viên / đồng nghiệp")
    mixed(d, [
        ("\"Omega của McDonald là phiên bản ", {}),
        ("hiện đại", {'bold': True}),
        (" của Cronbach alpha. Cả hai cùng đo ", {}),
        ("độ tin cậy nội tại", {'bold': True}),
        (" của một thang đo, ngưỡng đọc giống nhau (≥0,70 ổn, ≥0,80 tốt). Khác biệt là "
         "omega ", {}),
        ("không bắt buộc các câu hỏi đóng góp như nhau", {'bold': True}),
        (", nên thường ", {}),
        ("chính xác hơn alpha", {'bold': True}),
        (" khi thang đo có item nặng-nhẹ khác nhau — đó là lý do nhiều tạp chí "
         "đã yêu cầu báo cáo omega thay vì alpha.\"", {}),
    ])

    h(d, "Tham khảo")
    nr(d, "• McDonald, R. P. (1999). Test Theory: A Unified Treatment. Lawrence Erlbaum.", size=11)
    nr(d, "• Dunn, T. J., Baguley, T., & Brunsden, V. (2014). From alpha to omega: A practical "
          "solution to the pervasive problem of internal consistency estimation. "
          "British Journal of Psychology, 105(3), 399–412.", size=11)
    nr(d, "• Hayes, A. F., & Coutts, J. J. (2020). Use omega rather than Cronbach's alpha for "
          "estimating reliability. Communication Methods and Measures, 14(1), 1–24.", size=11)
    nr(d, "• Truy vấn RAG: tro-ly-nghien-cuu-tam-ly/rag_db_full (collection lo_au_full, "
          "47 chunks, embedding paraphrase-multilingual-MiniLM-L12-v2).", size=10, italic=True, color=GRAY)

    out = os.path.join(OUT_DIR, "McDonald_Omega_la_gi_cho_thay_25042026.docx")
    d.save(out)
    return out

# =====================================================================
# DOC 2 — Bài 8 (Wen 2020) đính chính giới tính & diễn giải OR
# =====================================================================
def make_doc_wen():
    d = new_doc()
    title(d, "BÀI 8 (WEN et al. 2020) — ĐÍNH CHÍNH CHIỀU GIỚI TÍNH\nVÀ CÁCH ĐỌC OR < 1")
    subtitle(d, "Trả lời câu hỏi của thầy — 25/04/2026")
    nr(d, "Trợ lý nghiên cứu — soạn từ truy vấn RAG (chunk QT08) + KG (kg_data) "
          "+ đối chiếu BANG_TOM_TAT_11_BAI_BAO.md.",
       size=10, italic=True, color=GRAY)

    h(d, "1. Câu hỏi của thầy")
    nr(d, "Trong bài 8, có câu sau làm thầy không hiểu là nam bị lo âu nặng hơn hay nữ:",
       italic=True)
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.7); p.paragraph_format.right_indent = Cm(0.7)
    r = p.add_run("\"NAM có nhiều khả năng phát triển lo âu trung bình và nặng HƠN NỮ "
                  "(OR lo âu trung bình = 0,649 cho nam, tức nữ cao hơn; OR lo âu nặng = 0,262 "
                  "cho nam, tức NỮ cao gấp gần 4 lần — ĐÂY LÀ TRÁI NGƯỢC: thực tế nữ > nam "
                  "trong nghiên cứu này). Áp lực học tập rất cao: OR = 11,579 cho lo âu nặng. "
                  "Hỗ trợ SKTT tại trường đầy đủ: OR = 0,562 (bảo vệ).\"")
    r.italic = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = GRAY

    h(d, "2. Trả lời thẳng")
    mixed(d, [
        ("Trong nghiên cứu Wen et al. (2020) — bài số 8 của tổng quan — ", {}),
        ("thực tế NỮ bị lo âu nặng hơn NAM", {'bold': True, 'color': RED}),
        (". Câu trong bản dịch bị diễn đạt ngược về mệnh đề mở đầu (\"Nam có nhiều khả năng "
         "phát triển… HƠN nữ\") nhưng phần trong ngoặc đã ghi đúng và có cảnh báo \"ĐÂY LÀ "
         "TRÁI NGƯỢC: thực tế nữ > nam\". Em đã đính chính lại để đọc liền mạch ở mục 5 dưới.", {}),
    ])

    h(d, "3. Cách đọc OR < 1 — quy tắc chung")
    nr(d, "OR (odds ratio) so sánh khả năng (odds) của một biến cố giữa 2 nhóm. Khi báo cáo "
          "OR cho biến \"nam\" (với nữ là nhóm tham chiếu):")
    nr(d, "• OR = 1,00  → nam và nữ có khả năng như nhau.")
    nr(d, "• OR > 1,00  → NAM có khả năng cao hơn nữ (gấp OR lần).")
    mixed(d, [
        ("• OR < 1,00  → ", {}),
        ("NAM có khả năng THẤP HƠN nữ", {'bold': True}),
        (" → suy ra ", {}),
        ("NỮ có khả năng cao hơn nam, gấp 1/OR lần", {'bold': True, 'color': RED}),
        (".", {}),
    ])

    h(d, "4. Áp dụng vào số liệu Wen 2020")
    tbl(d,
        ['Mức lo âu', 'OR cho nam (nữ = ref)', 'Cách đọc đúng', 'Nữ cao hơn nam'],
        [
            ['Lo âu trung bình', '0,649', 'Nam = 0,649 lần nữ', '1 / 0,649 ≈ 1,54 lần'],
            ['Lo âu nặng',       '0,262', 'Nam = 0,262 lần nữ', '1 / 0,262 ≈ 3,82 lần'],
        ],
        [3.8, 4.0, 4.5, 3.5])
    mixed(d, [
        ("Kết luận ngắn: ", {'bold': True}),
        ("ở mức lo âu trung bình nữ cao hơn nam ~1,5 lần; ở mức lo âu NẶNG nữ cao hơn nam "
         "gần ", {}),
        ("4 lần", {'bold': True, 'color': RED}),
        (" — tức là chiều giới tính trong Wen 2020 ", {}),
        ("phù hợp xu hướng chung", {'bold': True}),
        (" của hầu hết nghiên cứu (nữ > nam về lo âu nội hoá), không phải ngoại lệ.", {}),
    ])

    h(d, "5. Toàn cảnh các OR đáng chú ý của Wen 2020 (kèm diễn giải)")
    tbl(d,
        ['Yếu tố', 'OR', 'Diễn giải đầy đủ'],
        [
            ['Nam (vs nữ) — lo âu trung bình', '0,649',
             'Nữ cao hơn nam ~1,54 lần (yếu tố nguy cơ là nữ giới)'],
            ['Nam (vs nữ) — lo âu nặng', '0,262',
             'Nữ cao hơn nam ~3,82 lần (chênh lệch giới rõ rệt ở mức nặng)'],
            ['Áp lực học tập rất cao — lo âu nặng', '11,579',
             'Tăng nguy cơ ~11,6 lần — yếu tố nguy cơ MẠNH NHẤT trong bài'],
            ['Hỗ trợ SKTT tại trường đầy đủ', '0,562',
             'Yếu tố BẢO VỆ — giảm khoảng 44% nguy cơ lo âu nặng'],
        ],
        [4.5, 1.7, 8.3])

    h(d, "6. Vì sao bản tóm tắt cũ ghi sai \"Nam > nữ\"?")
    nr(d, "Có 2 lý do kỹ thuật khiến tóm tắt 11 bài lúc đầu (BANG_TOM_TAT_11_BAI_BAO.md) ghi "
          "ngược chiều giới tính cho Wen 2020:")
    nr(d, "(a) Wen 2020 dùng nữ làm nhóm tham chiếu khi báo OR cho nam, nên OR < 1 dễ bị đọc "
          "nhầm thành \"nam ít hơn ⇒ ghi 'nam nhiều hơn'\" do quên nghịch đảo.")
    nr(d, "(b) Bài dùng phương pháp LPA (Latent Profile Analysis) chia 3 hồ sơ "
          "(nhẹ 19,22%, trung bình 56,00%, nặng 24,78%) — % thô ở mỗi hồ sơ giữa nam và nữ "
          "có thể gần nhau, nhưng hồi quy logistic đa biến (sau khi kiểm soát covariates) "
          "mới cho thấy nữ vượt nam rõ rệt. Người tóm tắt ban đầu dựa vào % thô thay vì OR "
          "đã hiệu chỉnh.")
    mixed(d, [
        ("Lỗi này đã được ghi nhận ở dòng 134 file ", {}),
        ("BANG_TOM_TAT_11_BAI_BAO.md", {'italic': True}),
        (" với mục \"ĐÍNH CHÍNH (12/04/2026)\". Tuy nhiên các dòng 35 và 107 trong cùng file "
         "vẫn còn cụm \"Nam > nữ lo âu\" — em đề xuất sửa thành ", {}),
        ("\"Nữ > nam (OR cho nam = 0,649 trung bình; 0,262 nặng; p < 0,05)\"", {'bold': True}),
        (" để đồng bộ.", {}),
    ])

    h(d, "7. Đối chiếu với các bài khác trong tổng quan")
    nr(d, "Sau khi đính chính, Wen 2020 phù hợp với xu hướng \"nữ > nam\" của hầu hết bài "
          "trong tổng quan (Hoa 2024 Hà Nội, Ngô Anh Vinh 2024 DTTS, Vĩnh Lộc 2024, Islam "
          "2025 — 59 quốc gia LMIC: AOR nữ = 1,51, v.v.). Hai ngoại lệ THỰC SỰ là:")
    nr(d, "• Saikia et al. (2023) Assam, Ấn Độ: nam 30,0% vs nữ 18,9% (p = 0,049, n = 360) "
          "— do văn hoá đặc thù Đông Bắc Ấn.")
    nr(d, "• Xu et al. (2021) Trung Quốc: nam 10,11% vs nữ 9,66% (chênh 0,45 điểm %, n = "
          "373.216, COVID peak — clinically insignificant dù p < 0,001).")

    h(d, "8. Tham khảo & truy vết")
    nr(d, "• Bài gốc: Wen X, Lin Y, Liu Y, Starcevich K, Yuan F, Wang X, Xie X, Yuan Z (2020). "
          "A Latent Profile Analysis of Anxiety among Junior High School Students in Less "
          "Developed Rural Regions of China. Int J Environ Res Public Health, 17(11):4079. "
          "https://doi.org/10.3390/ijerph17114079", size=11)
    nr(d, "• Chunk RAG: tro-ly-nghien-cuu-tam-ly/rag_db_full → collection lo_au_full → "
          "paper_id = QT08 (key_finding: \"Áp lực OR=11,58; hỗ trợ trường OR=0,562\").",
       size=10, italic=True, color=GRAY)
    nr(d, "• KG node: 06_Scripts/kg_data/nodes.json → QT08 (n=900, location=TQ nông thôn, "
          "tool=DASS-21+LPA — lưu ý: bài gốc dùng MHT 100 mục, trường meta 'tool' trong KG ghi "
          "DASS-21+LPA cần kiểm chứng).", size=10, italic=True, color=GRAY)
    nr(d, "• Tóm tắt nội bộ: BANG_TOM_TAT_11_BAI_BAO.md, dòng 17 (thông tin bài), dòng 35 + "
          "107 (cần sửa), dòng 134 (đính chính 12/04/2026).", size=10, italic=True, color=GRAY)

    out = os.path.join(OUT_DIR, "Bai_8_Wen_2020_GioiTinh_dinh_chinh_cho_thay_25042026.docx")
    d.save(out)
    return out

if __name__ == '__main__':
    p1 = make_doc_omega()
    print('Wrote:', p1)
    p2 = make_doc_wen()
    print('Wrote:', p2)
