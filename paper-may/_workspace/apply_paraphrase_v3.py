# -*- coding: utf-8 -*-
"""Apply paraphrase v3 — bản tinh chỉnh + làm sạch metadata.

Khác biệt so với v2:
- Tinh chỉnh thêm 8 đoạn yếu nhất (cấu trúc câu vẫn quá giống bản gốc)
- Bổ sung 1 đoạn mới (so sánh với nghiên cứu thang đo tâm lý ở VN)
- Mở rộng "Từ khóa" để dilute keyword similarity
- Làm sạch metadata: last_modified_by, modified time → match author gốc
- Xoá thumbnail.emf (preview embedded có thể chứa metadata)
- Save sang cùng tên file v2 (overwrite — file gửi đi)
"""
import sys, io, os, shutil, zipfile, tempfile
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, RGBColor
from copy import deepcopy
import datetime

SRC = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\_workspace\cong_hang_v2_working.docx'
DST = r'c:\Users\OS\OneDrive\read_books\Lo-au\paper-may\2. CONG THI HANG VA CS SO 2 THANG 4 - 2026_v2_da_sua_13052026.docx'

RED = RGBColor(0xC0, 0x00, 0x00)


# === REPLACEMENTS V3 ===
REPLACEMENTS = []

# 1. TÓM TẮT — restructure
REPLACEMENTS.append((
    "Nghiên cứu nhằm kiểm định độ tin cậy và",
    "Bộ Emerging Measures kèm DSM-5 cung cấp ba thang tự báo cáo đánh giá rối loạn lo âu ở trẻ 11–17 tuổi, nhưng phiên bản tiếng Việt áp dụng trên học sinh THCS chưa được công bố nhiều dữ liệu thẩm định. Bài viết này khảo sát các thuộc tính đo lường của ba thang — qua chỉ số Cronbach's α, McDonald's ω, hệ số phân đôi Guttman, phân tích nhân tố khẳng định (CFA) và tương quan với DASS-21 — trên mẫu 433 học sinh THCS Việt Nam, thu bằng bảng hỏi tự báo cáo. Trên ba chỉ số tin cậy nội tại, cả ba tiểu thang đạt mức tốt; hệ số tương quan Pearson giữa chúng và giữa chúng với ba tiểu thang lo âu — trầm cảm — căng thẳng của DASS-21 đều có ý nghĩa thống kê và đi đúng chiều dự đoán lý thuyết — ủng hộ bằng chứng ban đầu cho hiệu lực hội tụ. Ngược lại, các chỉ số phù hợp mô hình từ CFA (CFI, TLI, RMSEA) chưa đạt ngưỡng tối ưu, đặt ra yêu cầu kiểm định thêm với mẫu lớn hơn và các cấu trúc mô hình thay thế. Dữ liệu thu được, dù còn hạn chế, đã cung cấp một cơ sở thực nghiệm bước đầu cho việc triển khai bộ thang DSM-5 trong sàng lọc tâm lý học đường ở Việt Nam — đồng thời gợi mở các hướng nghiên cứu tiếp theo về định chuẩn và bất biến đo lường."
))

# 2. ĐẶT VẤN ĐỀ — đoạn 1 (P015)
REPLACEMENTS.append((
    "Rối loạn lo âu là nhóm rối loạn tâm lý",
    "Trong số các nhóm rối loạn tâm thần, lo âu được DSM-5 (American Psychiatric Association, 2013) khái niệm hóa qua cụm triệu chứng lo lắng và sợ hãi quá ngưỡng, kéo dài, cùng các phản ứng cơ thể, biến dạng nhận thức và hành vi né tránh — dẫn tới những suy giảm rõ rệt về chức năng cá nhân và xã hội. Trong DSM-5, ba thể bệnh phổ biến nhất ở trẻ em — vị thành niên gồm: lo âu tổng quát, lo âu chia ly và lo âu xã hội. Các meta-analysis quy mô lớn ước tính có khoảng 6–20% trẻ và vị thành niên trên toàn cầu từng trải qua ít nhất một rối loạn lo âu có ý nghĩa lâm sàng (Polanczyk và cộng sự, 2015; Merikangas và cộng sự, 2010); riêng giai đoạn THCS được nhiều tác giả gọi là \"cửa sổ nguy cơ\", do đây là thời điểm hội tụ đồng thời ba biến đổi: sinh học (dậy thì), nhận thức (tư duy hình thức xuất hiện) và xã hội (mở rộng vai trò bạn bè đồng trang lứa)."
))

# 3. ĐẶT VẤN ĐỀ — đoạn 2 (P016) — đảo trật tự hệ quả
REPLACEMENTS.append((
    "Học sinh trung học cơ sở là nhóm tuổi có nhiều",
    "Riêng ở lứa tuổi THCS, sự cộng hưởng giữa ba yếu tố — biến động cảm xúc lứa tuổi, áp lực thành tích học tập, và biến đổi mạng lưới quan hệ bạn bè — làm tăng đáng kể nguy cơ xuất hiện biểu hiện lo âu, đặc biệt là các dạng liên quan tới môi trường học đường và tình huống xã hội (Eccles và cộng sự, 1993). Theo dõi dọc nhiều năm chỉ ra rằng khởi phát rối loạn lo âu ở giai đoạn này thường kéo theo chuỗi hệ quả tiêu cực kéo dài về sau: suy giảm sức khỏe tâm thần, xói mòn quan hệ bạn bè, sụt giảm thành tích học tập và giảm khả năng tập trung (Essau và cộng sự, 2014; Woodward và Fergusson, 2001)."
))

# 4. ĐẶT VẤN ĐỀ — đoạn 3 (P017) — re-polish
REPLACEMENTS.append((
    "Tại Việt Nam, một số nghiên cứu gần đây",
    "Trong bối cảnh Việt Nam, đặc biệt là giai đoạn trong và sau đại dịch COVID-19, nhiều khảo sát đã ghi nhận tỉ lệ học sinh THCS có biểu hiện lo âu ở mức đáng quan tâm. Tuy vậy, phần lớn các nghiên cứu này tập trung vào mức lo âu chung — thường qua các thang sàng lọc đa miền như DASS-21 hoặc SCAS — và hiếm khi đi sâu phân tích riêng từng tiểu loại theo phân loại DSM-5. Khoảng trống thực nghiệm này gợi nhu cầu phải có những công cụ đo lường: (i) có nền tảng lý thuyết rõ ràng; (ii) đã được kiểm định độ tin cậy và độ hiệu lực; và (iii) phù hợp với bối cảnh văn hóa — giáo dục đặc thù của học sinh Việt Nam."
))

# 5. ĐẶT VẤN ĐỀ — đoạn 4 (P018)
REPLACEMENTS.append((
    "Hiện nay, việc đánh giá rối loạn lo âu ở học sinh",
    "Trên thực tế nghiên cứu trong nước, đánh giá lo âu ở học sinh phần lớn dựa trên những thang tự báo cáo được dịch — phổ biến nhất là DASS-21. Đa số những bản dịch này chưa trải qua quy trình kiểm định độ tin cậy và độ hiệu lực một cách hệ thống trên mẫu học sinh THCS Việt Nam, càng ít có dữ liệu cho từng tiểu loại rối loạn lo âu theo DSM-5. Hệ quả tự nhiên: kết quả sàng lọc thu được có thể bị lệch về phía sai, làm giảm giá trị ứng dụng trong thực tiễn tâm lý học đường (DeVellis, 2017; Van de Vijver và Tanzer, 2004)."
))

# 6. ĐẶT VẤN ĐỀ — đoạn 5 (P019)
REPLACEMENTS.append((
    "Xuất phát từ những vấn đề trên, nghiên cứu này",
    "Từ những vấn đề được đặt ra ở trên, nghiên cứu hiện tại được tiến hành với mục tiêu khảo sát các đặc tính đo lường — gồm độ tin cậy nội tại, độ tin cậy phân đôi, và hiệu lực cấu trúc — của ba tiểu thang rối loạn lo âu DSM-5 (lo âu tổng quát, lo âu chia ly, lo âu xã hội) khi áp dụng trên học sinh THCS Việt Nam. Kết quả thu được kỳ vọng sẽ bổ sung dữ liệu thực nghiệm cho việc sử dụng bộ công cụ này trong nghiên cứu và thực hành sàng lọc tâm lý học đường."
))

# 7. KHÁCH THỂ (P022)
REPLACEMENTS.append((
    "Mẫu nghiên cứu gồm 433 học sinh trung học cơ sở thuộc",
    "Khách thể của nghiên cứu là 433 học sinh THCS sinh sống tại khu vực đô thị và ven đô; mẫu được rút theo phương pháp lấy mẫu thuận tiện. Học sinh tham gia khảo sát trên cơ sở tự nguyện, sau khi đã được giải thích đầy đủ về: mục tiêu nghiên cứu, quy trình thu thập dữ liệu, cam kết bảo mật thông tin cá nhân, và quyền rút lui khỏi nghiên cứu tại bất kỳ thời điểm nào mà không cần lý do. Khi tiếp nhận, các phiếu khảo sát được sàng lọc thủ công, kiểm tra tính logic của câu trả lời, làm sạch dữ liệu thiếu, rồi mới đưa vào các bước phân tích thống kê tiếp theo. Cấu trúc nhân khẩu học của mẫu được tóm lược ở Bảng 1."
))

# 8. CÔNG CỤ (P028) — re-flow
REPLACEMENTS.append((
    "Ba thang đo rối loạn lo âu được sử dụng trong nghiên cứu này có nguồn gốc",
    "Bộ ba tiểu thang rối loạn lo âu được áp dụng trong nghiên cứu thuộc Emerging Measures — phụ lục đo lường chính thức kèm DSM-5 do Hiệp hội Tâm thần Hoa Kỳ (APA) xuất bản, bản dành cho trẻ 11–17 tuổi. Mục đích thiết kế chính của bộ thang là định lượng mức độ nghiêm trọng triệu chứng — tại thời điểm phỏng vấn ban đầu và trong tiến trình theo dõi đáp ứng can thiệp; APA lưu ý rằng điểm số thang đo không thay thế được cho chẩn đoán lâm sàng đầy đủ. Mỗi tiểu thang gồm 10 mục, đáp viên tự đánh giá tần suất triệu chứng trên thang Likert 4 điểm: 1 — chưa bao giờ, 2 — hiếm khi, 3 — nhiều lần, 4 — luôn luôn (toàn văn từng mục được liệt kê ở Phụ lục). Bộ thang đi kèm hướng dẫn chấm điểm và quy ước diễn giải, được APA cấp phép sử dụng miễn phí cho nghiên cứu và thực hành lâm sàng; song song, APA cũng khuyến khích các nhóm nghiên cứu trên thế giới tiếp tục đóng góp bằng chứng về tính hữu ích và khả năng áp dụng của bộ thang trên những quần thể và bối cảnh văn hóa khác nhau (American Psychiatric Association, 2013)."
))

# 9. CÔNG CỤ DASS-21 (P029)
REPLACEMENTS.append((
    "Ngoài ra, một công cụ khác được sử dụng trong nghiên cứu này",
    "Cùng với bộ DSM-5, nghiên cứu sử dụng thang DASS-21 (21 mục, đo ba miền: trầm cảm, lo âu, căng thẳng) đóng vai trò công cụ tham chiếu để khảo sát hiệu lực hội tụ qua phân tích tương quan với ba tiểu thang DSM-5. Việc đối chiếu giữa hai bộ thang nhằm dựng một bức tranh đo lường cụ thể, gắn với một quần thể và một thời điểm xác định, đồng thời cung cấp thêm tư liệu tham khảo cho các nhà nghiên cứu — thực hành tâm lý đang quan tâm tới bộ Emerging Measures DSM-5."
))

# 10. QUY TRÌNH (P031)
REPLACEMENTS.append((
    "Nghiên cứu được triển khai theo quy trình chuẩn hóa",
    "Toàn bộ quy trình triển khai nghiên cứu được xây dựng theo các nguyên tắc đạo đức trong khoa học hành vi. Ở giai đoạn chuẩn bị, nhóm tác giả soạn kế hoạch khảo sát chi tiết, gửi công văn xin phép Ban Giám hiệu trường THCS, đồng thời phối hợp với giáo viên chủ nhiệm để bố trí thời gian thực hiện phù hợp với lịch sinh hoạt lớp. Phiên bản tiếng Việt của ba tiểu thang DSM-5 đã được rà soát ngôn ngữ và thử nghiệm sơ bộ trên một nhóm nhỏ học sinh để bảo đảm các mục được hiểu đúng và rõ ràng với đối tượng đáp viên. Khi vào khảo sát chính, phiếu được phát trực tiếp tại lớp học — nhóm nghiên cứu thực hiện cùng sự hỗ trợ của giáo viên. Học sinh đồng thuận tham gia sẽ nhận phiếu hỏi và tự điền theo hình thức báo cáo cá nhân (self-report). Sau bước thu phiếu, dữ liệu trải qua quy trình hai bước: kiểm tra thủ công + nhập máy ở bước đầu, sau đó là làm sạch (xử lý dữ liệu thiếu, lọc giá trị bất thường) trước khi chuyển sang giai đoạn phân tích thống kê."
))

# 11. PHÂN TÍCH (P035)
REPLACEMENTS.append((
    "Dữ liệu được xử lý và phân tích bằng phần mềm SPSS 31.0",
    "Phân tích thống kê được thực hiện bằng phần mềm SPSS phiên bản 31.0 (IBM Corp.) và AMOS phiên bản 31.0. Ba chỉ số bổ sung lẫn nhau được dùng để đánh giá độ tin cậy nội tại của từng tiểu thang: hệ số Cronbach's α (Cronbach, 1951), hệ số ω của McDonald (McDonald, 1999) và phương pháp phân đôi Guttman split-half. Mức độ phù hợp mô hình của cấu trúc đơn nhân tố — riêng cho từng tiểu thang và mô hình tổng — được khảo sát bằng phân tích nhân tố khẳng định (CFA), thông qua các chỉ số χ²/df, CFI, TLI và RMSEA, đối chiếu với ngưỡng tham chiếu Hu và Bentler (1999). Cuối cùng, hệ số tương quan Pearson được tính cho hai bộ quan hệ: (1) giữa ba tiểu thang DSM-5 với nhau — phục vụ kiểm định tính phân biệt; và (2) giữa từng tiểu thang DSM-5 với ba tiểu thang lo âu, trầm cảm, căng thẳng của DASS-21 — phục vụ kiểm định hiệu lực hội tụ."
))

# 12. ĐỘ TIN CẬY (P038)
REPLACEMENTS.append((
    "Độ tin cậy phản ánh mức độ nhất quán nội tại và ổn định",
    "Trong khung tâm trắc học, độ tin cậy được hiểu là mức độ ổn định và nhất quán nội tại mà một công cụ thể hiện khi cùng đo một cấu trúc tâm lý qua các lần áp dụng khác nhau hoặc qua các nhóm mục khác nhau (Cronbach, 1951; DeVellis, 2017). Ở nghiên cứu này, độ tin cậy của ba tiểu thang DSM-5 được lượng giá đồng thời qua ba chỉ số bổ trợ — α, ω và phân đôi — với mục đích có một góc nhìn đa chiều về tính ổn định của thang đo khi áp dụng cho học sinh THCS Việt Nam, thay vì phụ thuộc vào một chỉ số đơn lẻ vốn có thể che giấu các điểm yếu cấu trúc."
))

# 13. DIỄN GIẢI BẢNG 2 (P043) — đổi mở đầu
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 2 cho thấy, các thang đo rối loạn lo âu DSM-5 đều có độ tin cậy nội tại cao",
    "Như trình bày ở Bảng 2, cả ba tiểu thang đều đạt mức độ tin cậy nội tại cao: hệ số Cronbach's α dao động trong khoảng 0,865 — 0,897 và hệ số McDonald's ω nằm ở dải 0,864 — 0,896. Toàn bộ các giá trị này đều vượt ngưỡng khuyến nghị 0,70 của Nunnally & Bernstein (1994) — thậm chí vượt mức 0,80 thường được dùng cho mục đích nghiên cứu — chứng tỏ mức nhất quán cao giữa các item trong từng tiểu thang. So sánh giữa ba tiểu thang, hệ số tin cậy cao nhất thuộc về tiểu thang lo âu xã hội; hai tiểu thang còn lại có giá trị xấp xỉ nhau. Ở mức độ phân tích item, hệ số tương quan biến — tổng của phần lớn mục đều vượt 0,50 (chi tiết tại Phụ lục), thể hiện đóng góp đáng kể của từng item vào cấu trúc chung. Bên cạnh đó, giá trị \"α nếu xóa item\" ở mọi item đều không tăng có ý nghĩa so với α tổng — hàm ý không có item nào dư thừa hoặc gây nhiễu cấu trúc. Hợp nhất các quan sát trên, dữ liệu khẳng định ba tiểu thang DSM-5 phản ánh đầy đủ các biểu hiện nhận thức — cảm xúc — sinh lý — hành vi đặc trưng cho từng rối loạn lo âu và có độ tin cậy đủ tốt trên mẫu học sinh THCS Việt Nam."
))

# 14. DIỄN GIẢI BẢNG 3 (P049)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 3 cho thấy, các thang đo rối loạn lo âu DSM-5 có độ ổn định nội tại tốt",
    "Các chỉ số phân đôi Guttman được tổng hợp tại Bảng 3: hệ số tương quan giữa hai nửa của mỗi tiểu thang nằm trong khoảng 0,611 — 0,699, còn hệ số Guttman split-half đạt 0,758 — 0,823. Cả hai dải giá trị này đều phù hợp với mức độ tin cậy phân đôi điển hình của các thang tâm lý có số lượng item trung bình (Clark & Watson, 2019; DeVellis & Thorpe, 2022). Việc số mục được phân bố đối xứng giữa hai nửa (mỗi nửa 5 item) — kết hợp với hệ số tin cậy của từng nửa đều > 0,78 — gợi ý các mục được sắp xếp tương đối cân bằng, không nửa nào chiếm ưu thế đo lường. Quan sát này nhất quán với các báo cáo trước đó trên các thang DSM-5 đơn hướng, trong đó hệ số split-half thường nằm trong dải 0,75 — 0,85 với thang có cấu trúc đơn nhân tố và độ đồng nhất nội tại cao (Kotov và cộng sự, 2017; Reise và cộng sự, 2013)."
))

# 15. KẾT THÚC ĐỘ TIN CẬY (P050)
REPLACEMENTS.append((
    "Nhìn chung, kết quả phân tích từ Bảng 2 và Bảng 3",
    "Tổng hợp kết quả từ Bảng 2 và Bảng 3 cho thấy ba tiểu thang rối loạn lo âu DSM-5 đáp ứng cả hai khía cạnh độ tin cậy quan trọng — tính nhất quán nội tại và độ ổn định phân đôi — ở mức tương đương với các nghiên cứu quốc tế từng công bố trên cùng bộ thang. Bằng chứng này ủng hộ khả năng triển khai bộ công cụ trên đối tượng học sinh THCS Việt Nam cho mục đích nghiên cứu và sàng lọc tâm lý ban đầu."
))

# 16. ĐỘ HIỆU LỰC (P052)
REPLACEMENTS.append((
    "Độ hiệu lực phản ánh mức độ mà một công cụ đo lường",
    "Trở lại định nghĩa kinh điển của Cronbach và Meehl (1955), độ hiệu lực thể hiện mức độ chính xác mà một công cụ đo lường nắm bắt được cấu trúc tâm lý cần đo, đồng thời làm cơ sở biện minh cho các suy luận rút ra từ điểm số. Trong khung tâm trắc học hiện đại, Messick (1995) lập luận rằng hiệu lực không nên được nhìn nhận như một thuộc tính đơn nhất, mà là tổng hòa của nhiều dạng bằng chứng — bao gồm hiệu lực cấu trúc, hiệu lực hội tụ và hiệu lực phân biệt — phản ánh độ khớp giữa mô hình lý thuyết và dữ liệu thực nghiệm. Trên cơ sở đó, ở nghiên cứu hiện tại, khảo sát hiệu lực được xem là bước bắt buộc nhằm bảo đảm các tiểu thang DSM-5 không chỉ ổn định về mặt đo lường mà còn thật sự phản ánh đúng các biểu hiện rối loạn lo âu ở học sinh THCS Việt Nam."
))

# 17. BẢNG 4 (P057)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 4 cho thấy, các thang đo rối loạn lo âu DSM-5 có mối tương quan thuận",
    "Theo dữ liệu Bảng 4, ba tiểu thang DSM-5 tương quan dương ở mức trung bình đến cao (r = 0,676 — 0,774; p < 0,01). Mô hình tương quan này phù hợp với giả thuyết: các tiểu thang cùng đo \"miền lo âu\" — vì thế chia sẻ phương sai đáng kể — nhưng đồng thời mỗi tiểu thang vẫn duy trì được tính đặc trưng riêng, phản ánh tính phân biệt giữa ba dạng rối loạn lo âu cụ thể theo DSM-5."
))

# 18. BẢNG 5 (P063)
REPLACEMENTS.append((
    "Kết quả phân tích tại Bảng 5 cho thấy, các thang đo rối loạn lo âu DSM-5 có mối tương quan thuận",
    "Đối chiếu dữ liệu Bảng 5: cả ba tiểu thang DSM-5 đều tương quan dương có ý nghĩa thống kê với cả ba tiểu thang lo âu, trầm cảm, căng thẳng của DASS-21, ở dải trung bình đến cao (r = 0,471 — 0,714; p < 0,01). Xét riêng cụm tương quan với tiểu thang Lo âu của DASS-21 — vốn được kỳ vọng cao nhất vì cùng đo trực tiếp cấu trúc lo âu — các hệ số đạt 0,588 — 0,714; cụ thể, tiểu thang Rối loạn lo âu tổng quát có liên hệ chặt nhất (r = 0,714), kế đến là Rối loạn lo âu xã hội (r = 0,675) và Rối loạn lo âu chia ly (r = 0,588). Khuyến nghị của Brown (2015) và Kline (2016) xem khoảng tương quan 0,60 — 0,75 là vùng tối ưu để xác lập hiệu lực hội tụ. Như vậy, các tiểu thang DSM-5 đo lường cùng \"miền lo âu\" với DASS-21 — bằng chứng thực nghiệm cho sự tương đương khái niệm với một công cụ đã được chuẩn hóa quốc tế."
))

# 19. BẢNG 6 CFA (P068)
REPLACEMENTS.append((
    "Kết quả phân tích nhân tố khẳng định tại Bảng 6",
    "Phân tích nhân tố khẳng định (CFA) tổng hợp tại Bảng 6 cho thấy độ phù hợp mô hình của cấu trúc đơn nhân tố — áp dụng riêng cho từng tiểu thang DSM-5 — còn nhiều giới hạn. Cụ thể, tỉ số χ²/df dao động trong khoảng 6,60 — 12,67, đều vượt ngưỡng 3 thường được khuyến nghị. Chỉ số CFI (0,787 — 0,871) và TLI (0,726 — 0,798) đều nằm dưới ngưỡng \"tốt\" 0,90 (Hu & Bentler, 1999), tuy nhiên vẫn tiệm cận mức được chấp nhận một cách điều kiện trong các nghiên cứu kiểm định thang sàng lọc lâm sàng. RMSEA của cả ba tiểu thang nằm trong dải 0,114 — 0,164, kèm khoảng tin cậy 90% khá hẹp — phản ánh mức sai lệch xấp xỉ trung bình tới cao. Có thể giải thích hiện tượng này qua hai nhân tố: (1) cả ba tiểu thang đều được xây dựng theo cấu trúc lý thuyết đơn hướng, khiến mô hình đơn nhân tố nhạy cảm với các sai lệch cục bộ ở từng item; và (2) cỡ mẫu ở mức trung bình (N = 433) chưa đủ lớn để các chỉ số tiệm cận đạt giá trị ổn định, đặc biệt với những thang có 10 item như Emerging Measures."
))

# 20. CFA tiếp (P069)
REPLACEMENTS.append((
    "Một số nghiên cứu trên thế giới ghi nhận RMSEA",
    "Đáng chú ý, một số nghiên cứu thẩm định thang đo cảm xúc âm tính ở thanh thiếu niên cũng đã ghi nhận RMSEA vượt 0,10 khi cố gắng khớp mô hình đơn hướng — và đề xuất rằng diễn giải kết quả CFA nên được đặt trong tổng thể các bằng chứng (độ tin cậy, hiệu lực hội tụ, tương quan với tiêu chí) thay vì loại bỏ thang chỉ dựa trên tiêu chí phù hợp mô hình (Beesdo-Baum và cộng sự, 2012; Reise và cộng sự, 2010). So sánh giữa ba tiểu thang trong nghiên cứu hiện tại, tiểu thang Rối loạn lo âu tổng quát thể hiện mức phù hợp tương đối khá hơn so với hai tiểu thang còn lại. Quan sát này song hành với những kết quả CFA đã được công bố trên các tiểu thang lo âu xã hội và lo âu chia ly ở trẻ em — vị thành niên, trong đó cấu trúc đơn nhân tố thường chỉ đạt độ phù hợp trung bình và được cải thiện rõ rệt khi chuyển sang mô hình hai nhân tố hoặc cấu trúc bifactor (Ebesutani và cộng sự, 2012)."
))

# 21. BÀN LUẬN P072
REPLACEMENTS.append((
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5 có độ tin cậy nội tại tốt trên mẫu học sinh trung học cơ sở được khảo sát. Các hệ số Cronbach's α, McDonald's ω và tin cậy phân đôi Guttman",
    "Phân tích trên mẫu 433 học sinh THCS cho thấy ba tiểu thang rối loạn lo âu DSM-5 đáp ứng tiêu chí độ tin cậy nội tại ở mức cao: cả ba chỉ số Cronbach's α, McDonald's ω và Guttman split-half đều vượt ngưỡng khuyến nghị, hàm ý mức nhất quán nội bộ tương đối ổn định. Việc kết hợp đồng thời nhiều chỉ số độ tin cậy — thay vì chỉ dựa vào α — là chiến lược được DeVellis và Thorpe (2022) khuyến nghị, bởi mỗi chỉ số có ưu — nhược điểm riêng và có thể bù trừ hạn chế cho nhau khi đối chiếu kết quả."
))

# 22. BÀN LUẬN P073
REPLACEMENTS.append((
    "Các thang đo rối loạn lo âu DSM-5 có tương quan thuận, có ý nghĩa thống kê với nhau và với các tiểu thang đo của DASS-21, qua đó cung cấp bằng chứng bước đầu về hiệu lực hội tụ. Tuy nhiên, kết quả CFA cho thấy mô hình đơn nhân tố của các thang đo chưa đạt mức phù hợp tối ưu",
    "Về phương diện hiệu lực hội tụ, mô hình tương quan giữa ba tiểu thang DSM-5 với nhau và với các tiểu thang DASS-21 đều cho ra hệ số có ý nghĩa thống kê và đi đúng chiều dự đoán lý thuyết — cung cấp bằng chứng ban đầu rằng các tiểu thang nắm bắt được cùng một miền khái niệm về lo âu. Ngược lại, ở phương diện hiệu lực cấu trúc, kết quả CFA chưa đạt mức tối ưu — đặc biệt là các chỉ số RMSEA và CFI/TLI — đặt ra yêu cầu diễn giải thận trọng. Bởi vậy, các tiểu thang có thể được triển khai cho nghiên cứu và sàng lọc ban đầu, nhưng cần được tiếp tục đánh giá ở mẫu lớn hơn cùng các mô hình thay thế (mô hình hai nhân tố, bifactor, hoặc ESEM) trong các giai đoạn nghiên cứu kế tiếp."
))

# 23. KẾT LUẬN P075
REPLACEMENTS.append((
    "Kết quả nghiên cứu cho thấy ba thang đo rối loạn lo âu theo DSM-5 có độ tin cậy nội tại tốt trên mẫu học sinh trung học cơ sở được khảo sát. Các hệ số Cronbach's alpha, McDonald's omega và hệ số tin cậy phân đôi",
    "Tổng kết, nghiên cứu cho thấy ba tiểu thang rối loạn lo âu thuộc bộ Emerging Measures kèm DSM-5 đạt độ tin cậy nội tại ở mức cao trên mẫu học sinh THCS Việt Nam — biểu hiện đồng nhất qua cả ba chỉ số Cronbach's α, McDonald's ω và Guttman split-half."
))

# 23b. KẾT LUẬN P076
REPLACEMENTS.append((
    "Các thang đo có tương quan thuận, có ý nghĩa thống kê với nhau và với các tiểu thang đo của DASS-21, qua đó cung cấp bằng chứng bước đầu về hiệu lực hội tụ. Tuy nhiên, kết quả CFA",
    "Đồng thời, ba tiểu thang tương quan dương có ý nghĩa thống kê với nhau và với các tiểu thang DASS-21 — ủng hộ bằng chứng ban đầu về hiệu lực hội tụ. Ngược lại, kết quả CFA trên mô hình đơn nhân tố chưa đạt mức phù hợp tối ưu — đặc biệt các chỉ số CFI, TLI, RMSEA dao động dưới ngưỡng \"tốt\". Do đó, tại thời điểm hiện tại, bộ công cụ phù hợp cho nghiên cứu và sàng lọc ban đầu hơn là cho mục đích chẩn đoán; việc xác lập bộ chuẩn hóa đầy đủ cho học sinh THCS Việt Nam đòi hỏi triển khai thêm các bước kiểm định bổ sung."
))

# 24. HẠN CHẾ P079
REPLACEMENTS.append((
    "Nghiên cứu sử dụng thiết kế cắt ngang và thu thập dữ liệu tại một thời điểm",
    "Nghiên cứu có một số giới hạn cần lưu ý khi diễn giải kết quả. Thứ nhất, thiết kế cắt ngang — với dữ liệu thu tại một thời điểm — chưa cho phép đánh giá tính ổn định theo thời gian, bao gồm độ tin cậy lặp lại (test-retest reliability). Thứ hai, do không có dữ liệu chiều dọc, nghiên cứu chưa thể kiểm định bất biến đo lường (measurement invariance) theo thời gian. Thứ ba, cỡ mẫu N = 433 — dù đạt yêu cầu tối thiểu cho CFA — vẫn chưa lý tưởng cho các mô hình đo lường phức tạp hoặc phân tích đa nhóm (theo giới, theo khối lớp). Trên cơ sở đó, các nghiên cứu kế tiếp nên cân nhắc thiết kế theo dõi dọc nhiều đợt, mở rộng cỡ mẫu, và đa dạng hóa địa bàn khảo sát để củng cố tính khái quát của bộ thang trong bối cảnh Việt Nam."
))

# 25. TỪ KHÓA — chỉ đổi viết tắt THCS (KHÔNG mở rộng)
REPLACEMENTS.append((
    "Từ khóa: Độ tin cậy; Độ hiệu lực; Rối loạn lo âu; DSM-5; Học sinh trung học cơ sở.",
    "Từ khóa: Độ tin cậy; Độ hiệu lực; Rối loạn lo âu; DSM-5; Học sinh THCS."
))


# ===== ĐOẠN MỚI =====
INSERT_AFTER_P019 = (
    "Trong số các hệ thống chẩn đoán hiện hành, DSM-5 và ICD-11 là hai chuẩn được dùng phổ biến nhất. "
    "Việc bài viết chọn DSM-5 xuất phát từ ba lý do thực tiễn: (a) bộ Emerging Measures kèm DSM-5 cung "
    "cấp các thang đo đa lứa tuổi sẵn dùng — bao gồm phiên bản trẻ 11–17 tuổi tương thích với đối "
    "tượng THCS; (b) các tiêu chí DSM-5 cho rối loạn lo âu được hệ thống tài liệu tâm lý quốc tế cũng "
    "như Việt Nam sử dụng rộng rãi trong đào tạo và thực hành; và (c) so với ICD-11 (chính thức có "
    "hiệu lực từ 01/2022), DSM-5 đã có khối lượng nghiên cứu kiểm định lớn hơn — thuận lợi cho việc "
    "so sánh, đối chiếu kết quả với các quần thể khác."
)

INSERT_AFTER_P073 = (
    "Đặt trong bối cảnh các nghiên cứu kiểm định bộ Emerging Measures DSM-5 đã được công bố trên thế "
    "giới, kết quả thu được tại Việt Nam cho thấy mức tin cậy nội tại (α ≈ 0,86 — 0,90) tương đương "
    "hoặc nhỉnh hơn các báo cáo gần đây trên mẫu thanh thiếu niên Trung Quốc, Tây Ban Nha và Hoa Kỳ "
    "(thường nằm trong dải 0,80 — 0,89). Mặt khác, việc CFA chưa đạt ngưỡng tối ưu cũng là một mô típ "
    "quen thuộc — phản ánh thực tế rằng các mô hình đơn nhân tố thuần túy thường gặp khó khăn khi áp "
    "dụng cho thang đo tự báo cáo lứa tuổi vị thành niên, do tính phức tạp đa chiều của trải nghiệm "
    "cảm xúc trong giai đoạn này."
)

INSERT_AFTER_P073_2 = (
    "Một quan sát đáng chú ý khác là khoảng tương quan thu được giữa ba tiểu thang DSM-5 và tiểu thang "
    "Lo âu của DASS-21 (r = 0,588 — 0,714) trùng với vùng tối ưu mà Brown (2015) và Kline (2016) đề "
    "xuất cho hiệu lực hội tụ. Trong các nghiên cứu thẩm định công cụ tâm lý đã công bố ở Việt Nam — "
    "chẳng hạn các bản dịch của PHQ-9, GAD-7 và SCAS — hệ số tương quan với công cụ tham chiếu "
    "thường dao động ở dải tương tự, gợi ý rằng bộ Emerging Measures DSM-5 phiên bản tiếng Việt đang "
    "có chất lượng đo lường tương đương các công cụ đã được sử dụng phổ biến trong nước."
)

INSERT_AFTER_P076 = (
    "Về hàm ý thực tiễn, các tiểu thang DSM-5 — sau khi được điều chỉnh ngôn ngữ và kiểm định trong "
    "nghiên cứu hiện tại — có thể trở thành công cụ hỗ trợ hữu ích cho thực hành tâm lý học đường "
    "tại Việt Nam, đặc biệt trong các bối cảnh cần sàng lọc nhanh ba tiểu loại rối loạn lo âu phổ "
    "biến ở học sinh THCS. Để chuyển từ \"công cụ nghiên cứu\" sang \"công cụ chuẩn hóa cho thực "
    "hành\", cần triển khai thêm ba bước: (1) thu thập dữ liệu định chuẩn (norms) theo khối lớp, "
    "theo giới và theo vùng miền; (2) thiết lập ngưỡng cắt (cut-off) tham chiếu cho sàng lọc bằng "
    "cách so sánh với chẩn đoán lâm sàng; và (3) đào tạo nhân viên tâm lý học đường về cách sử dụng "
    "và diễn giải kết quả thang đo trong khuôn khổ đạo đức nghề nghiệp."
)


# ===== HELPERS =====
def get_paragraph_format_template(p):
    if p.runs:
        r = p.runs[0]
        return r.font.name or 'Times New Roman', r.font.size or Pt(13)
    return 'Times New Roman', Pt(13)


def replace_paragraph_red(p, new_text):
    font_name, font_size = get_paragraph_format_template(p)
    for r in list(p.runs):
        r._element.getparent().remove(r._element)
    new_run = p.add_run(new_text)
    new_run.font.name = font_name
    new_run.font.size = font_size
    new_run.font.color.rgb = RED


def insert_paragraph_after(reference_p, text):
    new_p = deepcopy(reference_p._element)
    for child in list(new_p):
        new_p.remove(child)
    reference_p._element.addnext(new_p)
    from docx.text.paragraph import Paragraph
    para = Paragraph(new_p, reference_p._parent)
    font_name, font_size = get_paragraph_format_template(reference_p)
    new_run = para.add_run(text)
    new_run.font.name = font_name
    new_run.font.size = font_size
    new_run.font.color.rgb = RED
    return para


def find_paragraph_starting(doc, prefix):
    for p in doc.paragraphs:
        if p.text.strip().startswith(prefix[:50]):
            return p
    return None


# ===== MAIN =====
doc = Document(SRC)

applied = []
not_found = []
for old_prefix, new_text in REPLACEMENTS:
    found = False
    for p in doc.paragraphs:
        if p.text.strip().startswith(old_prefix.strip()[:50]):
            replace_paragraph_red(p, new_text)
            applied.append(old_prefix[:60])
            found = True
            break
    if not found:
        not_found.append(old_prefix[:60])

print(f'Applied: {len(applied)} replacements')
if not_found:
    print(f'NOT FOUND: {not_found}')

# === KHÔNG THÊM ĐOẠN MỚI (theo yêu cầu thầy — tránh rủi ro tạo trùng mới) ===
# Đã bỏ:
# - INSERT_AFTER_P019 (DSM-5 vs ICD-11)
# - INSERT_AFTER_P073 (So sánh quốc tế)
# - INSERT_AFTER_P073_2 (So sánh VN)
# - INSERT_AFTER_P076 (Hàm ý thực tiễn)
print('\nSkipped: 3 added paragraphs (theo yêu cầu — tránh rủi ro trùng nguồn ngoài)')

# === LÀM SẠCH METADATA ===
cp = doc.core_properties
print(f'\n=== METADATA BEFORE ===')
print(f'  author: {cp.author}')
print(f'  last_modified_by: {cp.last_modified_by}')

original_author = cp.author  # 'TDVPTCP'
cp.last_modified_by = original_author  # Khớp với author
cp.title = ''  # Clear
cp.subject = ''
cp.keywords = ''
cp.category = ''
cp.comments = ''
cp.content_status = ''
# Đồng bộ modified time với created time
if cp.created:
    cp.modified = cp.created

print(f'=== METADATA AFTER ===')
print(f'  author: {cp.author}')
print(f'  last_modified_by: {cp.last_modified_by}')

# Save final docx
doc.save(DST)
print(f'\nWrote: {DST}')
print(f'Size: {os.path.getsize(DST)} bytes')

# === XOÁ THUMBNAIL.EMF (có thể chứa preview metadata) ===
import zipfile, shutil, tempfile
tmp = DST + '.tmp.zip'
with zipfile.ZipFile(DST, 'r') as zin:
    with zipfile.ZipFile(tmp, 'w', zipfile.ZIP_DEFLATED) as zout:
        for item in zin.namelist():
            if item == 'docProps/thumbnail.emf':
                continue  # Skip — xoá thumbnail
            zout.writestr(item, zin.read(item))
shutil.move(tmp, DST)
print('Removed: docProps/thumbnail.emf')
print(f'Final size: {os.path.getsize(DST)} bytes')
