# -*- coding: utf-8 -*-
"""
Fix expert review issues:
1. tuoi teen -> vi thanh nien (consistency)
2. ga trong nam tinh -> tomboy/co gai nam tinh (mistranslation of "manly chick")
3. bộ co -> Bộ này co (capitalization)
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
PATH = os.path.join(ROOT, '03_Ban-dich', 'VN022_UNICEF_VN_2022_SchoolFactors_FULL.docx')

d = Document(PATH)
changes = 0

FIXES = [
    # Lỗi 1: tuổi teen (informal) → vị thành niên (formal, consistency)
    ('hiểu và hỗ trợ con tuổi teen', 'hiểu và hỗ trợ con cái vị thành niên'),
    ('tuổi teen của các em', 'vị thành niên của các em'),

    # Lỗi 2: "gà trống nam tính" mistranslation of "manly chick"
    # In EN: "manly chick" = LGBTQ slang for tomboy/masculine girl
    ('‘gà trống nam tính’', '"tomboy" (con gái có tính cách nam tính — nguyên bản: "manly chick")'),
    ('gà trống nam tính', 'tomboy (con gái có tính cách nam tính)'),

    # Lỗi 3: capitalize "bộ" when refer to specific Bộ
    ('chia sẻ rằng bộ có một chương trình', 'chia sẻ rằng Bộ Y tế có một chương trình'),
]

for p in d.paragraphs:
    if p.text.strip():
        new_text = p.text
        for find, rep in FIXES:
            if find in new_text:
                new_text = new_text.replace(find, rep)
        if new_text != p.text:
            for run in p.runs:
                run.text = ''
            if p.runs:
                p.runs[0].text = new_text
            changes += 1

for t in d.tables:
    for r in t.rows:
        for c in r.cells:
            for p in c.paragraphs:
                if p.text.strip():
                    new_text = p.text
                    for find, rep in FIXES:
                        if find in new_text:
                            new_text = new_text.replace(find, rep)
                    if new_text != p.text:
                        for run in p.runs:
                            run.text = ''
                        if p.runs:
                            p.runs[0].text = new_text
                        changes += 1

d.save(PATH)
print(f'Expert review fixes: {changes} changes')

# Verify
d2 = Document(PATH)
txt = '\n'.join(p.text for p in d2.paragraphs if p.text.strip())
print(f'\nVerification:')
print(f'  "tuổi teen" remaining: {txt.count("tuổi teen")}')
print(f'  "gà trống" remaining: {txt.count("gà trống")}')
print(f'  "tomboy" added: {txt.count("tomboy")}')
print(f'  "rằng bộ có": {txt.count("rằng bộ có")}')
print(f'  "rằng Bộ Y tế có": {txt.count("rằng Bộ Y tế có")}')
