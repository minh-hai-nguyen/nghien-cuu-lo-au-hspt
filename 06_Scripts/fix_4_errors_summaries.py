# -*- coding: utf-8 -*-
"""Fix 4 real errors found by QA check."""
import sys, os, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAGE_W = 16.0

# Helpers (same as before)
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)
def make_doc():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3); sec.right_margin = Cm(2)
    return doc
def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name='Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
def P(doc, text, bold=False, italic=False, size=12, color=None):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.name='Times New Roman'; r.font.size=Pt(size)
    r.bold=bold; r.italic=italic
    if color: r.font.color.rgb = color
def table(doc, headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
def save(doc, name):
    p = os.path.join(TT_DIR, name); doc.save(p)
    d = Document(p); chars = sum(len(x.text) for x in d.paragraphs)
    print(f'  FIXED {name}: {chars} chars, {len(d.tables)} tables')

# ============================================================
# FIX 1: VN25 Hải Phòng — sai tỷ lệ, thiếu tác giả, bối cảnh không đủ
# ============================================================
print('[FIX 1] VN25 Hải Phòng 2024')
doc = make_doc()
H(doc, 'Tóm tắt bài VN25 — Sức khoẻ tâm thần học sinh THPT + GDNN-GDTX huyện Vĩnh Bảo, Hải Phòng 2023', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, 'Công trình "Thực trạng sức khoẻ tâm thần và một số yếu tố liên quan của học sinh tại '
       '2 trường trung học phổ thông huyện Vĩnh Bảo, Hải Phòng năm 2023" của Phạm Thị Ngọc, '
       'Hoàng Thị Giang, Phạm Khánh Linh, Vũ Thị Châu (2024), đăng trên Tạp chí Y học Dự '
       'phòng Việt Nam, tập 34, số 1 Phụ bản — 2024, trang 115. Đơn vị chính: Trường Đại '
       'học Y Dược Hải Phòng. Khách thể: 420 học sinh tại 2 cơ sở giáo dục ở huyện Vĩnh Bảo, '
       'Hải Phòng — (1) THPT Cộng Hiền và (2) Trung tâm Giáo dục nghề nghiệp – Giáo dục '
       'thường xuyên (GDNN-GDTX) Vĩnh Bảo. Thời gian khảo sát: từ tháng 10/2022 đến tháng '
       '05/2023. Ngày nhận bài: 26/01/2024; Ngày đăng: 12/04/2024.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Thiết kế cắt ngang mô tả. Công cụ: DASS-21 phiên bản tiếng Việt đo cả 3 trục (trầm '
       'cảm, lo âu, căng thẳng) + bảng hỏi nhân khẩu học và các yếu tố nguy cơ (giới tính, '
       'quan hệ bạn bè, áp lực học tập và thi cử, kỳ vọng gia đình, mâu thuẫn gia đình, '
       'bạo lực học đường). Phân tích chi-square + hồi quy logistic đa biến.')

H(doc, 'Kết quả nghiên cứu', level=2)
table(doc,
    ['Rối loạn', 'THPT Cộng Hiền', 'GDNN-GDTX Vĩnh Bảo'],
    [
        ['Trầm cảm', '42,38 %', '32,38 %'],
        ['Lo âu', '53,81 %', '43,33 %'],
        ['Căng thẳng', '49,05 %', '28,57 %'],
    ],
    widths=[5.0, 5.0, 5.5])
P(doc, 'LƯU Ý — Tỷ lệ lo âu CAO HƠN đáng kể so với Hà Nội (Hoa 2024: 40,6 %) và Trung Quốc '
       '(Dong 2025: 41,4 %). Tỷ lệ ở THPT chính quy (Cộng Hiền 53,81 %) cao hơn ở GDNN-GDTX '
       '(Vĩnh Bảo 43,33 %), gợi ý áp lực học thuật đóng vai trò quan trọng.', italic=True)

P(doc, 'Các yếu tố nguy cơ có ý nghĩa thống kê:', bold=True)
P(doc, '• Trầm cảm: giới nữ, không có mối quan hệ hoà đồng với bạn bè.')
P(doc, '• Lo âu: áp lực cao với học tập và thi cử, không có mối quan hệ hoà đồng với bạn '
       'bè, bị bạo lực học đường.')
P(doc, '• Căng thẳng (riêng HS THPT Cộng Hiền): áp lực cao với kỳ vọng gia đình, gia đình '
       'mâu thuẫn.')

H(doc, 'Phản biện', level=2)
P(doc, 'Điểm mạnh: Mẫu 420 HS từ 2 cơ sở giáo dục — đặc biệt bao gồm cả GDNN-GDTX (đối '
       'tượng ít được nghiên cứu). Sử dụng DASS-21 đo cả 3 trục — phù hợp so sánh trực '
       'tiếp với Dong 2025 (TQ) và VN29 Duong 2025 TPHCM. Tỷ lệ lo âu ở Cộng Hiền (53,81 %) '
       'thuộc nhóm cao nhất VN, gần với Long An (57,2 %) và cao hơn Hà Nội (40,6 %). '
       'Phân biệt rõ THPT chính quy vs GDNN-GDTX là đóng góp quan trọng.')
P(doc, 'Hạn chế: Chỉ 1 huyện (Vĩnh Bảo) của Hải Phòng — không đại diện toàn thành phố. '
       'Mẫu thuận tiện tại 2 cơ sở, không ngẫu nhiên. Cắt ngang không kết luận nhân quả. '
       'Các yếu tố nguy cơ chỉ phân tích riêng cho HS Cộng Hiền ở phần căng thẳng — thiếu '
       'so sánh trực tiếp giữa 2 cơ sở.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Mở rộng ra nhiều huyện của Hải Phòng. So sánh đa thành phố (Hải Phòng vs Hà Nội '
       'vs TPHCM) bằng cùng DASS-21. NC dọc theo dõi cùng nhóm 2-3 năm. Phân tích sâu sự '
       'khác biệt giữa THPT chính quy và GDNN-GDTX để có can thiệp phù hợp từng nhóm.')

P(doc, 'Đánh giá chất lượng: ⭐⭐⭐ Trung bình–Khá. Tạp chí trong nước, mẫu vừa, DASS-21 '
       'tiếng Việt, phát hiện phân biệt 2 loại cơ sở có giá trị.', bold=True)
save(doc, 'VN25_HaiPhong_2024.docx')

# ============================================================
# FIX 2: QT44 Cai (not Cao) Resilience
# ============================================================
print('[FIX 2] QT44 Cai Resilience (sửa Cao → Cai)')
doc = make_doc()
H(doc, 'Tóm tắt bài QT44 — SR + MA can thiệp resilience tại trường', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"School-based interventions for resilience in children and adolescents: a '
       'systematic review and meta-analysis of randomized controlled trials" của '
       'Chenyi Cai, Zhengyang Mei, Zirui Wang, Shi Luo (2025), đăng trên Frontiers in '
       'Psychiatry vol. 16, article 1594658, tr. 1–15. DOI 10.3389/fpsyt.2025.1594658. '
       'Q1, IF ≈ 4,7. Open Access. Đơn vị: School of Physical Education, Southwest '
       'University, Chongqing, China (tác giả chính Chenyi Cai). Tổng quan hệ thống '
       'và phân tích tổng hợp các RCT về can thiệp khả năng phục hồi (resilience) tại '
       'trường cho trẻ em + vị thành niên. LƯU Ý: tác giả chính là Cai (không phải Cao) '
       'theo Pinyin của 蔡 (họ Cai).')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống các RCT trên PubMed, Web of Science, APA PsycINFO, Google '
       'Scholar về can thiệp resilience tại trường. Đánh giá chất lượng bằng Cochrane '
       'Risk of Bias tool (RoB 2) cho RCT. Random-effects meta-analysis. Sử dụng '
       'Standardized Mean Difference (SMD; Cohen d) với KTC 95 % để tổng hợp hiệu ứng.')
P(doc, 'Resilience = khả năng thích ứng tốt trước nghịch cảnh, căng thẳng hoặc sang chấn. '
       'Khác CBT (giảm triệu chứng) — tăng yếu tố BẢO VỆ: tự trọng, kỹ năng ứng phó, '
       'kết nối xã hội, lạc quan.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Các nghiên cứu được thực hiện tại nhiều quốc gia: Mỹ (n = 10), Trung Quốc (n = 9), '
       'Úc (n = 4), Pakistan (n = 2), Ấn Độ (n = 2) và các nước khác — bằng chứng từ '
       'nhiều văn hoá khác nhau.')
P(doc, 'Can thiệp resilience tại trường CÓ HIỆU QUẢ tăng resilience và giảm triệu chứng '
       'SKTT ở trẻ em/VTN. Tuy nhiên, kích thước hiệu ứng NHỎ–TRUNG BÌNH và heterogeneity '
       'CAO. Nghĩa là: hiệu quả tồn tại nhưng không lớn, và có sự khác biệt đáng kể giữa '
       'các nghiên cứu (do khác biệt thiết kế, mẫu, văn hoá).')

H(doc, 'Phản biện', level=2)
P(doc, 'Frontiers in Psychiatry Q1 IF ≈ 4,7, Open Access PMC. SR + MA RCTs — bằng chứng '
       'tốt. Resilience là yếu tố BẢO VỆ — bổ sung cho CBT, không thay thế. Phù hợp với '
       'Trần Thảo Vi 2025 (lạc quan trung gian, β = −0,24) và Ireland MyWorld 2024 '
       '(resilience + tự trọng quan trọng tăng theo thời gian). Đa dạng địa lý (Mỹ, '
       'Trung, Úc, Pakistan, Ấn) bao phủ nhiều văn hoá — có 9 RCT từ Trung Quốc nên '
       'bối cảnh Á Đông khá đầy đủ. Hạn chế: hiệu lực nhỏ–TB; heterogeneity cao; chưa '
       'có NC từ Việt Nam hoặc ĐNA.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Phát triển module resilience tích hợp vào can thiệp CBT trường VN. RCT VN cho '
       'VTN có yếu tố nguy cơ cao. Thiết kế đo cả triệu chứng (CBT) lẫn nguồn lực bảo '
       'vệ (resilience).')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SR+MA Q1, định hướng module bổ trợ.', bold=True)
save(doc, 'QT44_Cai_Resilience_MA_2025.docx')
# Remove old file with wrong surname
old = os.path.join(TT_DIR, 'QT44_Cao_Resilience_MA_2025.docx')
if os.path.exists(old):
    os.remove(old)
    print('  REMOVED old file: QT44_Cao_Resilience_MA_2025.docx')

# ============================================================
# FIX 3: QT45 Sasaki — sửa UMIN ID
# ============================================================
print('[FIX 3] QT45 Sasaki — sửa UMIN ID')
doc = make_doc()
H(doc, 'Tóm tắt bài QT45 — RCT iCBT đa trung tâm cho subthreshold SAD tại Nhật Bản', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for '
       'Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: Multicenter '
       'Randomized Controlled Trial" của Sasaki, N. và cộng sự (2024), đăng trên JMIR '
       'Pediatrics and Parenting vol. 7, e55786, tr. 1–13. DOI 10.2196/55786. UMIN-CTR '
       'UMIN000050064 (LƯU Ý: ID chính xác là 000050064, không phải 000049768). Đối '
       'tượng: 77 học sinh / sinh viên có triệu chứng SAD DƯỚI NGƯỠNG (subthreshold) '
       'tại 6 đại học và 1 trường THPT Nhật Bản, từ 12/2022 đến 10/2023.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'RCT đa trung tâm 2 nhánh: nhóm can thiệp (n = 38) nhận iCBT TỰ HỌC HOÀN TOÀN qua '
       'web, không có hỗ trợ người; nhóm đối chứng (n = 39) chờ đợi. Can thiệp 8 module, '
       'theo dõi đến tháng 3 sau khởi động. Đo lường: SIAS (lo âu xã hội chính), PHQ-9 '
       '(trầm cảm). Phân tích ANCOVA với trầm cảm là biến đồng. Fisher exact test cho '
       'response/recovery ratios.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Nhóm iCBT giảm có ý nghĩa thống kê triệu chứng lo âu xã hội so với nhóm đối chứng. '
       'Tỷ lệ đáp ứng và phục hồi:')
table(doc, ['Chỉ số', 'Nhóm can thiệp', 'OR vs chờ (95% CI)', 'p'],
      [['Đáp ứng (response)', '61 %', '4,97 (1,61–16,53)', '0,003'],
       ['Phục hồi (recovery)', '68 %', '3,95 (1,32–12,56)', '0,008'],
       ['Remission ratio', '—', '2,01', 'ns']],
      widths=[5.0, 3.5, 4.0, 3.0])

H(doc, 'Phản biện', level=2)
P(doc, 'RCT đa trung tâm — bằng chứng cao. Đặc biệt quan trọng vì SAD DƯỚI NGƯỠNG là dân '
       'số rất lớn và quan trọng — can thiệp ở giai đoạn này có thể NGĂN CHẶN tiến triển '
       'thành rối loạn đầy đủ ("indicated prevention"). Mô hình iCBT TỰ HỌC HOÀN TOÀN có '
       'chi phí gần bằng 0 sau khi phát triển nội dung — phù hợp triển khai đại trà cho '
       'VN. Mâu thuẫn với Walder 2025 (DMHI không hướng dẫn g chỉ ~ 0,3) — có thể do nội '
       'dung được thiết kế tốt hoặc hiệu ứng bị overestimate do mẫu nhỏ. Hạn chế: n = 77 '
       'nhỏ, 1 quốc gia, và KTC rộng (1,61–16,53 cho OR đáp ứng) cho thấy độ chính xác '
       'không cao.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'Adapt iCBT Nhật cho VN — sàng lọc HS subthreshold SAD bằng SIAS-17, can thiệp web '
       'tiếng Việt. So sánh có/không hỗ trợ người trong bối cảnh VN. RCT lớn hơn (n > 200) '
       'để xác nhận hiệu quả.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. RCT đa trung tâm Q2, mô hình prevention sớm khả thi.',
  bold=True)
save(doc, 'QT45_Sasaki_Japan_iCBT_2024.docx')

# ============================================================
# FIX 4: QT46 Academic Stress SR — bổ sung tác giả + số NC chính xác
# ============================================================
print('[FIX 4] QT46 Academic Stress SR — bổ sung tác giả')
doc = make_doc()
H(doc, 'Tóm tắt bài QT46 — Tổng quan hệ thống can thiệp căng thẳng học tập ở HS trung học', level=1)
P(doc, 'Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát', bold=True)
P(doc, '"Academic Stress Interventions in High Schools: A Systematic Literature Review" '
       'của Jagiello, T., Belcher, J., Neelakandan, A., Boyd, K. & Wuthrich, V.M. (2025), '
       'đăng trên Child Psychiatry and Human Development vol. 56, tr. 1836–1869 (Q1, IF '
       '≈ 3,3) — Springer. DOI 10.1007/s10578-024-01667-5. Online 4/3/2024, in print '
       '2025. Tổng quan hệ thống các chương trình can thiệp giảm/phòng ngừa CĂNG THẲNG '
       'HỌC TẬP riêng ở HS trung học trên toàn cầu. Bài dài 34 trang — là tổng quan hệ '
       'thống ĐẦU TIÊN tập trung riêng vào chương trình trường học cho stress học tập.')

H(doc, 'Phương pháp nghiên cứu', level=2)
P(doc, 'Tìm kiếm hệ thống các nghiên cứu đánh giá hiệu quả chương trình trường học trong '
       'giảm hoặc phòng ngừa stress học tập ở HS trung học. Tiêu chí lựa chọn nghiêm ngặt. '
       'Phân loại theo (a) loại can thiệp (CBT, mindfulness, thư giãn, kỹ năng quản lý, '
       'yoga, ...); (b) hình thức (universal phổ quát vs targeted có mục tiêu); (c) người '
       'cung cấp (chuyên gia vs giáo viên). Đánh giá chất lượng phương pháp.')

H(doc, 'Kết quả nghiên cứu', level=2)
P(doc, 'Tìm thấy 31 nghiên cứu đủ tiêu chuẩn từ 13 quốc gia. Phát hiện chính:', bold=True)
P(doc, '• Chất lượng phương pháp của đa số nghiên cứu KÉM (poor quality) — nhiều NC sử '
       'dụng nhóm đối chứng KHÔNG can thiệp.')
P(doc, '• Bằng chứng MẠNH NHẤT cho các chương trình dựa trên CBT — phù hợp với lý thuyết '
       'về căng thẳng học tập.')
P(doc, '• CẢ 2 mô hình universal VÀ targeted đều có thể có lợi — không có mô hình nào vượt '
       'trội rõ ràng.')
P(doc, '• Stress học tập = yếu tố MẠNH NHẤT ảnh hưởng SKTT VTN châu Á — đã được xác nhận '
       'bởi VN21 Trần Thảo Vi (β = 4,73 cho học thêm), Wen QT08 (OR = 11,58 cho áp lực HT), '
       'Norway QT21, UNICEF VN22 (47 % HS học thêm > 3h/tuần).')

table(doc, ['Loại can thiệp', 'Hiệu quả', 'Phù hợp VN'],
      [['CBT cho stress HT', 'BẰNG CHỨNG MẠNH NHẤT', 'Cần thích ứng VH'],
       ['Mindfulness/thiền', 'TB — kết quả lẫn lộn', 'UK 8.376 HS thất bại'],
       ['Quản lý thời gian', 'Hứa hẹn', 'Rất phù hợp VN'],
       ['Thư giãn cơ', 'TB', 'Dễ triển khai'],
       ['Yoga/hoạt động thể chất', 'Hứa hẹn', 'Phù hợp PE'],
       ['Kỹ năng ứng phó', 'TB–Khá', 'Kết hợp CBT']],
      widths=[5.0, 4.5, 6.0])

H(doc, 'Phản biện', level=2)
P(doc, 'Child Psychiatry and Human Development Q1 IF ≈ 3,3, Springer. 34 trang — rất toàn '
       'diện. Tổng quan ĐẦU TIÊN về can thiệp stress học tập riêng ở HS trung học (theo '
       'lời tuyên bố của tác giả). 31 NC từ 13 quốc gia — phạm vi địa lý tốt. CBT có bằng '
       'chứng mạnh nhất — phù hợp QT29 BMC NMA. Hạn chế: chất lượng phương pháp của đa số '
       'NC được đánh giá là KÉM; nhiều NC dùng nhóm chứng không can thiệp; không có '
       'meta-analysis định lượng; chưa rõ đại diện LMIC/châu Á.')

H(doc, 'Hướng nghiên cứu tiếp theo', level=2)
P(doc, 'RCT CHẤT LƯỢNG CAO can thiệp stress HT riêng tại VN — CBT + quản lý thời gian + '
       'giảm tải học thêm. So sánh trong bối cảnh học sinh VN (sức ép đa chiều: cha mẹ, '
       'thi đại học, học thêm). Đánh giá hiệu quả dài hạn 1–2 năm. Sử dụng nhóm đối chứng '
       'active (không phải waiting list thuần) để cải thiện chất lượng bằng chứng.')
P(doc, 'Đánh giá chất lượng: ⭐⭐⭐⭐ Cao. SLR đầu tiên, Q1 Springer, 31 NC từ 13 quốc gia, '
       'rất phù hợp đề tài.', bold=True)
save(doc, 'QT46_AcademicStress_SR_2025.docx')

print('\n=== ALL 4 ERRORS FIXED ===')
