# -*- coding: utf-8 -*-
"""Fix NBSP at run start (cross-run pattern).
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

d = Document(FILE)
fixed = 0
for i, p in enumerate(d.paragraphs):
    runs = p.runs
    for ri in range(len(runs)):
        curr = runs[ri]
        # If run starts with NBSP and previous run ends with punctuation : ; , ! ?
        if curr.text and curr.text[0] == '\xa0':
            if ri > 0 and runs[ri-1].text and runs[ri-1].text[-1] in ':;,!?':
                # Replace leading NBSP with space
                curr.text = ' ' + curr.text[1:]
                fixed += 1
            else:
                # NBSP at start without preceding punct - replace with regular space anyway
                curr.text = ' ' + curr.text[1:]
                fixed += 1

d.save(FILE)
print(f"Fixed {fixed} NBSP at run-start")

# Verify
d2 = Document(FILE)
remain = sum(1 for p in d2.paragraphs if re.search(r'[,;:!?]\xa0|\xa0[,;:!?]', p.text))
print(f"NBSP near punct remaining: {remain}")
remain_lead = 0
for p in d2.paragraphs:
    for r in p.runs:
        if r.text.startswith('\xa0'):
            remain_lead += 1
print(f"NBSP at run-start remaining: {remain_lead}")
