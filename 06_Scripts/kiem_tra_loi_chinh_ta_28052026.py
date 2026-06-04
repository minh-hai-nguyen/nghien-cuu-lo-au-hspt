# -*- coding: utf-8 -*-
"""Kiem tra loi chinh ta + trinh bay - KHONG SUA gi het.
Chi bao cao loi de NCS quyet dinh sua.
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

print("="*80)
print("BAO CAO LOI CHINH TA + TRINH BAY")
print("File: 01_LuanAn_v3_1_FixCoping_28052026.docx")
print("="*80)

# Helper: get paragraph context (chuong/section)
def show_context(idx):
    return f"para {idx}"

# ================================================================
# 1. DOUBLE SPACES IN TEXT
# ================================================================
print("\n[1] DOUBLE SPACES (khoang trang doi) trong text")
double_space_paras = []
for i, p in enumerate(d.paragraphs):
    if '  ' in p.text:  # 2+ spaces
        double_space_paras.append((i, p.text.strip()[:120]))
print(f"  Tim thay {len(double_space_paras)} doan co khoang trang doi")
for i, txt in double_space_paras[:10]:
    print(f"    para {i}: {txt}")
if len(double_space_paras) > 10:
    print(f"    ... va {len(double_space_paras)-10} doan khac")


# ================================================================
# 2. SPACE BEFORE PUNCTUATION (Vietnamese standard: no space before . , ; : ! ?)
# ================================================================
print("\n[2] DAU CACH TRUOC DAU CAU (sai chuan: 'abc ,' / 'abc .' / 'abc ;')")
space_before_punct = []
for i, p in enumerate(d.paragraphs):
    matches = re.findall(r'\w\s+[.,;:!?]', p.text)
    if matches:
        # Get first instance
        m = re.search(r'\w\s+[.,;:!?]', p.text)
        if m:
            pos = m.start()
            ctx = p.text[max(0,pos-30):pos+30]
            space_before_punct.append((i, ctx, len(matches)))
print(f"  Tim thay {len(space_before_punct)} doan")
for i, ctx, n in space_before_punct[:10]:
    print(f"    para {i} ({n}x): ...{ctx}...")
if len(space_before_punct) > 10:
    print(f"    ... va {len(space_before_punct)-10} doan khac")


# ================================================================
# 3. MISSING SPACE AFTER PUNCTUATION
# ================================================================
print("\n[3] THIEU DAU CACH SAU DAU CAU (sai chuan: 'abc,xyz' khong space)")
no_space_after_punct = []
for i, p in enumerate(d.paragraphs):
    # Find punctuation followed by alphanumeric (no space)
    # Exclude common patterns: decimal "1.5", "1,5" (Vietnamese decimal), "1-2", "a.m."
    matches = re.findall(r'[a-zA-Zà-ỹÀ-Ỹ][.,;:!?][a-zA-Zà-ỹÀ-Ỹ]', p.text)
    if matches:
        no_space_after_punct.append((i, matches[:3], p.text[:120]))
print(f"  Tim thay {len(no_space_after_punct)} doan")
for i, matches, txt in no_space_after_punct[:10]:
    print(f"    para {i}: {matches} | {txt}")
if len(no_space_after_punct) > 10:
    print(f"    ... va {len(no_space_after_punct)-10} doan khac")


# ================================================================
# 4. INCONSISTENT NUMBERING (e.g., '3.4.2.1' appears twice)
# ================================================================
print("\n[4] SO THU TU TIEU DE BI TRUNG LAP (vi du 3.4.2.1 xuat hien 2 lan)")
heading_nums = {}
import re
for i, p in enumerate(d.paragraphs):
    if 'Heading' not in p.style.name and 'Tiêu đề' not in p.style.name:
        continue
    text = p.text.strip()
    m = re.match(r'^(\d+(?:\.\d+)*)', text)
    if m:
        num = m.group(1)
        if num not in heading_nums:
            heading_nums[num] = []
        heading_nums[num].append((i, text[:80]))
dupes = {k: v for k, v in heading_nums.items() if len(v) > 1}
print(f"  Tim thay {len(dupes)} so trung lap:")
for num, paras in dupes.items():
    print(f"  '{num}' xuat hien {len(paras)} lan:")
    for idx, txt in paras:
        print(f"    para {idx}: {txt}")


# ================================================================
# 5. STRAIGHT QUOTES vs CURLY QUOTES (Vietnamese standard prefers curly)
# ================================================================
print("\n[5] DAU NGOAC THANG (\") vs CONG (“”)")
straight_quotes = 0
curly_quotes = 0
for p in d.paragraphs:
    straight_quotes += p.text.count('"')
    curly_quotes += p.text.count('“') + p.text.count('”')
print(f"  Straight \" : {straight_quotes}")
print(f"  Curly  “ ”: {curly_quotes}")
if straight_quotes > 0 and curly_quotes > 0:
    print(f"  CO TRON LAN - nen thong nhat 1 kieu")


# ================================================================
# 6. INCONSISTENT TERMS (vi du 'sức khoẻ' vs 'sức khỏe')
# ================================================================
print("\n[6] THUAT NGU KHONG NHAT QUAN (chinh ta giua cac bien the)")
inconsistent_pairs = [
    ('sức khoẻ', 'sức khỏe'),  # standard: sức khỏe
    ('hoà', 'hòa'),
    ('thuý', 'thúy'),
    ('thuỷ', 'thủy'),
    ('huỳnh', 'huỷnh'),  # standard: huỳnh
    ('cảm xúc', 'cảm súc'),
    ('rốc', 'rốc'),
    ('Tâm lý', 'Tâm Lý'),
    ('Tâm thần', 'Tâm Thần'),
    ('Áp lực', 'áp lực'),
]
for a, b in inconsistent_pairs:
    count_a = 0; count_b = 0
    for p in d.paragraphs:
        count_a += p.text.count(a)
        count_b += p.text.count(b)
    if count_a > 0 and count_b > 0:
        print(f"  '{a}' ({count_a}x) vs '{b}' ({count_b}x)")


# ================================================================
# 7. MULTIPLE CONSECUTIVE BLANK PARAGRAPHS
# ================================================================
print("\n[7] NHIEU DOAN TRONG LIEN TIEP (gay khoang trang lon)")
consecutive_empty = 0
runs_of_empty = []
prev_empty = False
empty_run_start = None
for i, p in enumerate(d.paragraphs):
    is_empty = not p.text.strip()
    if is_empty:
        if not prev_empty:
            empty_run_start = i
        consecutive_empty += 1
    else:
        if consecutive_empty >= 3 and empty_run_start is not None:
            runs_of_empty.append((empty_run_start, i-1, consecutive_empty))
        consecutive_empty = 0
        empty_run_start = None
    prev_empty = is_empty
print(f"  Tim thay {len(runs_of_empty)} cum >= 3 doan trong:")
for start, end, n in runs_of_empty[:10]:
    print(f"    para {start}-{end}: {n} doan trong")


# ================================================================
# 8. PARENTHESES BALANCE (mat tro)
# ================================================================
print("\n[8] DAU NGOAC KHONG CAN BANG (so '(' khac '(' hoac mat ')' )")
unbalanced = []
for i, p in enumerate(d.paragraphs):
    opens = p.text.count('(')
    closes = p.text.count(')')
    if opens != closes:
        unbalanced.append((i, opens, closes, p.text[:100]))
print(f"  Tim thay {len(unbalanced)} doan khong can bang")
for i, o, c, txt in unbalanced[:10]:
    print(f"    para {i}: ({o} vs )({c}) | {txt}")


# ================================================================
# 9. MIXED ENGLISH/VIETNAMESE IN ABBREV
# ================================================================
print("\n[9] VIET TAT KHONG NHAT QUAN (RLLA vs RLLALT, etc.)")
abbrev_check = ['RLLALT', 'RLLA', 'RLLATQ', 'LATQ', 'LACL', 'LAXH', 'YTNC', 'YTBV',
                'ALHT', 'NDT', 'BNHD', 'GBT', 'LTT', 'HTCM', 'HTBB', 'BPDP',
                'THCS', 'THPT', 'NCS', 'GAD', 'DASS', 'CBT', 'PHQ']
for ab in abbrev_check:
    count = 0
    for p in d.paragraphs:
        # Word boundary check
        count += len(re.findall(rf'\b{re.escape(ab)}\b', p.text))
    if count > 0:
        print(f"  {ab}: {count} occurrences")


# ================================================================
# 10. NUMBERED LIST INCONSISTENCY
# ================================================================
print("\n[10] DANH SACH SO THU TU KHONG NHAT QUAN")
# Check headings only - look for jumps in numbering
print("  (kiem tra so chuong tu CHƯƠNG 1, 2, 3 - so the heading sequences)")
all_headings_with_num = []
for i, p in enumerate(d.paragraphs):
    if 'Heading' not in p.style.name and 'Tiêu đề' not in p.style.name:
        continue
    text = p.text.strip()
    m = re.match(r'^(\d+(?:\.\d+)*)', text)
    if m:
        all_headings_with_num.append((i, m.group(1), text[:70]))

# Sequence check: each "x.y.z" should appear after "x.y.(z-1)" or first of "x.y.1"
# Simplistic: list all H2-H4 sequences per parent
groups = {}
for i, num, txt in all_headings_with_num:
    parts = num.split('.')
    if len(parts) >= 2:
        parent = '.'.join(parts[:-1])
        groups.setdefault(parent, []).append((int(parts[-1]), i, num, txt))

issues_found = 0
for parent, children in groups.items():
    children.sort()
    seq = [c[0] for c in children]
    expected = list(range(1, len(seq) + 1))
    if seq != expected:
        # Look for duplicates
        seen = set(); dupes = []
        for s in seq:
            if s in seen: dupes.append(s)
            seen.add(s)
        if dupes:
            issues_found += 1
            print(f"  '{parent}.X': sequence {seq} (expected {expected}) - DUP {dupes}")
            for n, idx, num, txt in children:
                print(f"    para {idx}: {num} - {txt}")
        elif len(seq) > 0 and seq[0] != 1:
            issues_found += 1
            print(f"  '{parent}.X': bat dau tu {seq[0]} (nen bat dau 1)")
        elif any(seq[i] - seq[i-1] != 1 for i in range(1, len(seq))):
            issues_found += 1
            print(f"  '{parent}.X': sequence khong lien tuc {seq}")
print(f"  Total issues: {issues_found}")


print()
print("="*80)
print("HET BAO CAO")
print("="*80)
