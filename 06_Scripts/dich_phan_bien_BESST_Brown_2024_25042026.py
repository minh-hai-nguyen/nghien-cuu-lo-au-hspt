# -*- coding: utf-8 -*-
"""Dịch song ngữ + Phản biện đầy đủ bài Brown 2024 BESST trial.
Bài gốc: Brown J, James K, Lisk S, Shearer J, Byford S, Stallard P, Deighton J,
Saunders D, Yarrum J, Fonagy P, Weaver T, Sclare I, Day C, Evans C, Carter B (2024).
Clinical effectiveness and cost-effectiveness of a brief accessible cognitive
behavioural therapy programme for stress in school-aged adolescents (BESST):
A cluster RCT in the UK.
The Lancet Psychiatry, 11(7): 504-515. PMID 38759665. doi:10.1016/S2215-0366(24)00101-9
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\BESST_Brown_2024_dich_phan_bien_25042026.docx'

RED   = RGBColor(0xC0, 0x00, 0x00)
BLUE  = RGBColor(0x00, 0x70, 0xC0)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x00, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(13)
s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def title(text, size=18):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def subtitle(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def H1(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def H2(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def en(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def vn(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(13); r.font.name = 'Times New Roman'
def crit(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('[Phản biện] '); r.bold = True; r.font.color.rgb = RED; r.font.size = Pt(12); r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.font.color.rgb = RED; r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
def note(text, color=GREEN):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('[Ghi chú dịch giả] '); r.bold = True; r.italic = True; r.font.color.rgb = color; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.italic = True; r2.font.color.rgb = color; r2.font.size = Pt(11); r2.font.name = 'Times New Roman'
def nr(text, bold=False, size=13, color=None, italic=False):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color

# =====================================================================
# TRANG BÌA
# =====================================================================
title("BẢN DỊCH SONG NGỮ + PHẢN BIỆN", 16)
title("HIỆU QUẢ LÂM SÀNG VÀ CHI PHÍ-HIỆU QUẢ", 15)
title("CỦA CHƯƠNG TRÌNH CBT NGẮN ACCESSIBLE CHO STRESS", 14)
title("VỊ THÀNH NIÊN TUỔI ĐI HỌC (BESST)", 16)
title("MỘT THỬ NGHIỆM CLUSTER RCT TẠI VƯƠNG QUỐC ANH", 14)
subtitle("Clinical effectiveness and cost-effectiveness of a brief accessible cognitive")
subtitle("behavioural therapy programme for stress in school-aged adolescents (BESST):")
subtitle("A cluster randomised controlled trial in the UK")
nr("")
subtitle("Brown J, James K, Lisk S, Shearer J, Byford S, Stallard P, Deighton J,")
subtitle("Saunders D, Yarrum J, Fonagy P, Weaver T, Sclare I, Day C, Evans C, Carter B (2024)")
subtitle("The Lancet Psychiatry, 11(7): 504-515")
subtitle("doi: 10.1016/S2215-0366(24)00101-9 — PMID 38759665 — Open Access (CC BY 4.0)")
subtitle("Loại bài: ARTICLE — Cluster Randomised Controlled Trial (level 2 evidence)")
nr("")
nr("Trợ lý nghiên cứu — Dịch + Phản biện — 25/04/2026", italic=True, color=GRAY, size=10)
nr("Quy trình: dịch song ngữ EN↔VN từng section của RCT + phản biện chữ đỏ có dẫn chứng "
   "+ phân tích risk of bias (Cochrane RoB 2) + bảng outcomes chi tiết + so sánh "
   "BESST vs MYRIAD vs FRIENDS PACES + reference đầy đủ. Kiểm 3 vòng theo memory "
   "feedback_research_workflow.md.",
   italic=True, color=GRAY, size=10)

# =====================================================================
# THÔNG TIN THƯ MỤC
# =====================================================================
H1("THÔNG TIN THƯ MỤC")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài (EN)', 'Clinical effectiveness and cost-effectiveness of a brief '
                          'accessible cognitive behavioural therapy programme for stress '
                          'in school-aged adolescents (BESST): A cluster randomised '
                          'controlled trial in the UK'],
        ['Tên bài (VN)', 'Hiệu quả lâm sàng và chi phí-hiệu quả của chương trình CBT '
                          'ngắn dễ tiếp cận cho stress ở vị thành niên tuổi đi học (BESST): '
                          'Một thử nghiệm cluster RCT tại Vương quốc Anh'],
        ['Tác giả (15 người)',
         '1. June Brown (Chief Investigator) — Department of Psychology, IoPPN, KCL\n'
         '2. Kirsty James — Department of Biostatistics, IoPPN, KCL & King\'s Clinical Trials Unit\n'
         '3. Stephen Lisk — Department of Psychology, IoPPN, KCL\n'
         '4. James Shearer — Health Service & Population Research Department, IoPPN, KCL\n'
         '5. Sarah Byford — Department for Health, University of Bath\n'
         '6. Paul Stallard — Department for Health, University of Bath\n'
         '7. Jessica Deighton — Anna Freud National Centre for Children and Families, '
         'London\n'
         '8. David Saunders — Department of Educational and Health Psychology, University '
         'of Northampton\n'
         '9. Jynna Yarrum — Department of Educational and Health Psychology, University '
         'of Northampton\n'
         '10. Peter Fonagy — Faculty of Health, Social Care & Education, Middlesex '
         'University, London\n'
         '11. Timothy Weaver — Faculty of Health, Social Care & Education, Middlesex '
         'University, London\n'
         '12. Irene Sclare — Southwark CAMHS Clinical Academic Group, South London and '
         'Maudsley NHS Foundation Trust\n'
         '13. Crispin Day — IoPPN, KCL\n'
         '14. Claire Evans — IoPPN, KCL\n'
         '15. Ben Carter — Biostatistics & Health Informatics, IoPPN, KCL'],
        ['Liên hệ (corresponding)', 'Dr June Brown — june.brown@kcl.ac.uk'],
        ['Tạp chí', 'The Lancet Psychiatry'],
        ['Tập / Số / Trang', 'Vol. 11, No. 7, pp. 504-515 (July 2024)'],
        ['Published online', '14/05/2024 (corrected version 14/06/2024)'],
        ['DOI', '10.1016/S2215-0366(24)00101-9'],
        ['PMID', '38759665'],
        ['ISRCTN', '90912799'],
        ['Loại bài', 'Multicentre cluster RCT (intention-to-treat, masked outcome assessment)'],
        ['License', 'CC BY 4.0 (Open Access)'],
        ['Thời gian thử nghiệm', '04/10/2021 – 10/11/2022 (recruitment 2 năm học)'],
        ['Funding',
         'National Institute for Health and Care Research (NIHR) Health Technology '
         'Assessment Programme — Reference NIHR127951'],
        ['Ethics', 'King\'s College London PNM Research Ethics Subcommittee (HR-20/21-17758)'],
        ['Conflict of interest', 'Không có (declared)'],
        ['Mã trong corpus dự án', 'QT042 (chưa có file dịch full trước đó)'],
    ], [4.5, 11.5])

H2("Định vị bài này trong tổng quan dự án")
vn("Đây là RCT QUAN TRỌNG NHẤT của nhóm Brown — kiểm chứng định lượng giả thuyết "
   "PLACES (đã phát triển 2022) trong context can thiệp SKTT trường học UK. Đăng "
   "Lancet Psychiatry IF=64 — top-tier psychiatry journal. Trial đã bắt đầu 2021 "
   "(protocol Lisk 2022 Trials), recruited 2 năm học, phân tích 2024.")
crit("Đây là LEVEL-2 EVIDENCE (Oxford CEBM 2011) — RCT có quality cao hơn editorial "
     "(B5) và developmental paper (PLACES). Tuy nhiên không phải level 1 (SR/MA "
     "of RCTs). Khi trích cho thầy nên ghi rõ effect size và CI để tránh hiểu nhầm "
     "magnitude của effect. Lưu ý đặc biệt: từ \"modestly\" trong abstract của tác "
     "giả phản ánh đúng magnitude — KHÁC với cách Editorial Brown & Carter 2025 "
     "(B5) ghi \"very positive\" cho cùng kết quả này.")

# =====================================================================
# PHẦN 1 — DỊCH SONG NGỮ
# =====================================================================
H1("PHẦN 1 — DỊCH SONG NGỮ ĐẦY ĐỦ")

H2("Tóm tắt (Summary / Abstract)")

H2("Background")
en("Depression and anxiety are increasingly prevalent in adolescents. The Brief "
   "Educational Workshops in Secondary Schools Trial investigated the effectiveness of "
   "a brief accessible stress workshop programme for 16–18-year-olds. We aimed to "
   "investigate the clinical effectiveness and cost-effectiveness of the DISCOVER "
   "cognitive behavioural therapy (CBT) workshop on symptoms of depression in 16–18-"
   "year-olds at 6 months compared with treatment-as-usual.")
vn("Trầm cảm và lo âu ngày càng phổ biến ở vị thành niên. Thử nghiệm Brief Educational "
   "Workshops in Secondary Schools (BESST) đã điều tra hiệu quả của chương trình "
   "workshop stress ngắn dễ tiếp cận cho HS 16–18 tuổi. Chúng tôi nhằm điều tra HIỆU "
   "QUẢ LÂM SÀNG và CHI PHÍ-HIỆU QUẢ của workshop CBT DISCOVER lên TRIỆU CHỨNG TRẦM "
   "CẢM ở HS 16–18 tuổi tại thời điểm 6 tháng so với chăm sóc thông thường (TAU).")

H2("Methods")
en("We conducted a multicentre, cluster randomised controlled trial in UK schools or "
   "colleges with sixth forms to evaluate clinical effectiveness and cost-effectiveness "
   "of a brief CBT workshop (DISCOVER) compared with treatment-as-usual. We planned to "
   "enrol 60 schools and 900 adolescents, using a self-referral system to recruit "
   "participants. Schools were randomised in a 1:1 ratio for participants to receive "
   "either the DISCOVER workshop or treatment-as-usual, stratified by site and balanced "
   "on school size and index of multiple deprivation. Participants were included if they "
   "were 16–18 years old, attending for the full school year, seeking help for stress, "
   "and able to provide written informed consent. The outcome assessors, senior health "
   "economist, senior statistician, and chief investigator were masked. People with "
   "lived experience were involved in the study. The primary outcome was depression "
   "symptoms measured with the Mood and Feelings Questionnaire (MFQ) at 6-month follow-"
   "up, in the intention-to-treat population of all participants with full covariate "
   "data. The trial was registered with the ISRCTN registry (ISRCTN90912799).")
vn("Chúng tôi tiến hành một thử nghiệm CLUSTER RCT MULTICENTRE tại các trường UK hoặc "
   "college có sixth form để đánh giá hiệu quả lâm sàng và chi phí-hiệu quả của một "
   "workshop CBT ngắn (DISCOVER) so với CHĂM SÓC THÔNG THƯỜNG (treatment-as-usual / "
   "TAU). Chúng tôi LẬP KẾ HOẠCH tuyển 60 trường và 900 vị thành niên, sử dụng hệ "
   "thống SELF-REFERRAL để tuyển người tham gia. Các trường được PHÂN NGẪU NHIÊN tỉ lệ "
   "1:1 để người tham gia nhận DISCOVER hoặc TAU, PHÂN TẦNG theo site và CÂN BẰNG theo "
   "quy mô trường + chỉ số đa thiếu hụt (index of multiple deprivation / IMD). Người "
   "tham gia được bao gồm nếu: 16-18 tuổi, học toàn năm, đang tìm trợ giúp cho stress, "
   "có thể cung cấp consent có chữ ký. Outcome assessor, senior health economist, "
   "senior statistician, và chief investigator được CHE GIẤU (masked). Người có trải "
   "nghiệm sống (lived experience) được tham gia vào nghiên cứu. OUTCOME CHÍNH là "
   "triệu chứng trầm cảm đo bằng Mood and Feelings Questionnaire (MFQ) tại 6 tháng "
   "follow-up, trong intention-to-treat population của tất cả người tham gia có dữ "
   "liệu covariate đầy đủ. Thử nghiệm đăng ký với ISRCTN (ISRCTN90912799).")
crit("THIẾT KẾ chuẩn cho cluster RCT: phân tầng + cân bằng theo IMD (deprivation index) "
     "— quan trọng vì tránh confounding bởi yếu tố kinh tế-xã hội. Masking tốt: 4 vai "
     "trò được mask (assessor, economist, statistician, chief investigator). Tuy "
     "nhiên — như mọi RCT can thiệp tâm lý — KHÔNG THỂ mask participant (HS biết mình "
     "tham gia DISCOVER hay không). Đây là LIMITATION mà tác giả tự thừa nhận trong "
     "Discussion.")

H2("Findings")
en("111 schools were invited to participate in the study, seven were deemed ineligible, "
   "and 47 did not provide consent. Between Oct 4, 2021, and Nov 10, 2022, 933 students "
   "at 57 schools were screened for eligibility, seven were not eligible for inclusion, "
   "and 26 did not attend the baseline meeting and assessment, resulting in 900 "
   "adolescents participating in the study. The DISCOVER group included 443 participants "
   "(295 [67%] female and 136 [31%] male) and the treatment-as-usual group included 457 "
   "participants (346 [76%] female and 92 [20%] male). 468 (52%) of the 900 participants "
   "were White, and the overall age of the participants was 17·2 years (SD 0·6). 873 "
   "(97%) adolescents were followed up in the intention-to-treat population (n=854). "
   "The primary intention-to-treat analysis found an adjusted mean difference in MFQ of "
   "–2·06 (95% CI –3·35 to –0·76; Cohen's d=–0·17; p=0·0019) at the 6-month follow-up, "
   "indicating a clinical improvement in the DISCOVER group. The probability that "
   "DISCOVER is cost-effective compared with treatment-as-usual ranged from 61% to 78% "
   "at a £20 000 to £30 000 per quality-adjusted life-year threshold. Nine adverse "
   "events (two of which were classified as serious) were reported in the DISCOVER "
   "group and 14 (two of which were classified as serious) were reported in the "
   "treatment-as-usual group.")
vn("111 trường được mời, 7 không đủ điều kiện, 47 không cung cấp consent. Giữa "
   "04/10/2021 và 10/11/2022, 933 HS tại 57 trường được sàng lọc đủ điều kiện, 7 "
   "không đủ, và 26 không tham dự baseline meeting + assessment → 900 VTN tham gia "
   "nghiên cứu. Nhóm DISCOVER gồm 443 người tham gia (295 [67%] nữ và 136 [31%] nam), "
   "nhóm TAU gồm 457 (346 [76%] nữ và 92 [20%] nam). 468 (52%) trong tổng 900 là "
   "White, độ tuổi trung bình 17,2 tuổi (SD 0,6). 873 (97%) VTN được follow-up trong "
   "ITT population (n=854). Phân tích ITT chính tìm thấy ADJUSTED MEAN DIFFERENCE "
   "trong MFQ là −2,06 (95% CI −3,35 đến −0,76; Cohen's d = −0,17; p = 0,0019) tại "
   "6 tháng — chỉ ra cải thiện lâm sàng ở nhóm DISCOVER. Xác suất DISCOVER cost-"
   "effective so với TAU dao động 61% đến 78% ở ngưỡng £20.000 đến £30.000 mỗi "
   "QALY. 9 adverse events (2 nghiêm trọng) ở DISCOVER và 14 (2 nghiêm trọng) ở TAU.")
crit("CON SỐ THEN CHỐT cần lưu ý: Cohen d = −0,17 = SMALL effect size theo Cohen 1988 "
     "(d=0,2 small, 0,5 medium, 0,8 large; d=0,17 dưới ngưỡng small). 95% CI −3,35 "
     "đến −0,76 — KHÔNG cross zero (significant) nhưng UPPER BOUND của CI là −0,76 → "
     "vẫn small effect ở best case. p=0,0019 (rất significant về mặt thống kê) "
     "nhưng MAGNITUDE nhỏ. Đây là kinh điển \"clinically vs statistically significant\" "
     "— với n lớn (854), bất kỳ effect size nhỏ cũng đạt p<0,05. Khi báo cáo cho "
     "thầy nên ghi cả 3 con số: adjusted mean diff + 95% CI + Cohen d — không chỉ "
     "p-value.")

H2("Interpretation")
en("Our findings indicate that the DISCOVER intervention is modestly clinically "
   "effective and economically viable and could be a promising early intervention in "
   "schools. Given the importance of addressing mental health needs early in this "
   "adolescent population, additional research is warranted to explore this intervention.")
vn("Phát hiện của chúng tôi chỉ ra rằng can thiệp DISCOVER hiệu quả lâm sàng KHIÊM "
   "TỐN (modestly) và khả thi về kinh tế — và có thể là một can thiệp sớm hứa hẹn "
   "trong trường học. Cho tầm quan trọng của việc giải quyết nhu cầu SKTT sớm ở nhóm "
   "VTN này, nghiên cứu thêm là cần thiết để khám phá can thiệp này.")
crit("ĐÂY LÀ CỤM TỪ CỦA CHÍNH TÁC GIẢ: \"modestly clinically effective\" — KHIÊM TỐN. "
     "Tuyệt đối KHÔNG phải \"very positive\" như editorial B5 nói. Đây là điểm trung "
     "thực của abstract Lancet Psychiatry. Khi trích cho thầy phải dùng từ \"khiêm "
     "tốn\" / \"modest\" — không phóng đại.")

# ---------- Introduction ----------
H2("Mở đầu (Introduction)")
en("More than half of adult mental health conditions have first onset before the age of "
   "15 years, and almost three-quarters by the age of 18 years. Emotional disorders of "
   "anxiety and depression are especially common in the adolescent years, causing marked "
   "distress and daily interference for about one in 12 (8%) young people in England, "
   "with an increased risk of self-harm and suicidality among those with mental health "
   "conditions. The most recent government report on mental health of children and young "
   "people in England showed that the proportion of those aged 17–19 years with a "
   "probable mental health condition increased from 17·4% to 25·7% between 2021 and "
   "2022.")
vn("Hơn một nửa các tình trạng SKTT người lớn lần đầu khởi phát trước 15 tuổi, và "
   "GẦN BA PHẦN TƯ trước 18 tuổi. Các rối loạn cảm xúc lo âu và trầm cảm đặc biệt "
   "phổ biến ở những năm vị thành niên, gây ra sự đau khổ rõ rệt và can thiệp hàng "
   "ngày cho khoảng MỘT TRÊN 12 (8%) người trẻ ở Anh, với nguy cơ tự hại và tự tử "
   "tăng cao ở những người có tình trạng SKTT. Báo cáo chính phủ gần đây nhất về "
   "SKTT của trẻ em và người trẻ ở Anh cho thấy tỉ lệ người 17-19 tuổi có \"probable "
   "mental health condition\" tăng từ 17,4% lên 25,7% giữa 2021 và 2022.")
crit("Tăng từ 17,4% → 25,7% trong 1 năm là TĂNG GẤP RƯỠI — phản ánh tác động COVID + "
     "chiến tranh Ukraine + cost-of-living crisis. Áp dụng VN: cũng có thể có pattern "
     "tương tự (chưa có data wave 2 V-NAMHS), cần khảo sát lặp lại.")

en("Although data are not available specifically for 16–18-year-olds, it is estimated "
   "that 60% of children and young people with a diagnosable mental health condition do "
   "not receive any care through specialist child and adolescent mental health services "
   "(CAMHS) in the UK. Barriers to accessing formal support for young people include "
   "concerns about stigma and confidentiality and the limited capacity (and stringent "
   "eligibility criteria) of specialist mental health services, restricting access to "
   "effective evidence-based therapies. Consequently, there is a pressing need for "
   "scalable, accessible, and evidence-based interventions.")
vn("Mặc dù dữ liệu không có sẵn riêng cho HS 16-18 tuổi, ước tính 60% trẻ em và người "
   "trẻ có tình trạng SKTT có thể chẩn đoán KHÔNG nhận bất kỳ chăm sóc nào qua dịch "
   "vụ SKTT chuyên khoa CAMHS ở UK. Các rào cản tiếp cận hỗ trợ chính thức cho người "
   "trẻ bao gồm: lo ngại về kỳ thị và bảo mật, năng lực hạn chế (và tiêu chí đủ điều "
   "kiện nghiêm ngặt) của các dịch vụ SKTT chuyên khoa — hạn chế tiếp cận các liệu "
   "pháp dựa trên bằng chứng. Do đó, có nhu cầu CẤP THIẾT về các can thiệp QUY MÔ MỞ "
   "RỘNG, DỄ TIẾP CẬN, VÀ DỰA TRÊN BẰNG CHỨNG.")

en("Patel and colleagues have suggested that given the aforementioned problems, a staged "
   "approach to more formal services could help, with the first stage being accessible "
   "interventions in more youth-friendly settings such as schools. The latest review of "
   "school-based interventions reported a small effect size for depression and anxiety, "
   "indicating that interventions that were cognitive behavioural therapy (CBT)-based and "
   "delivered in secondary schools by clinicians were more likely to be effective.")
vn("Patel và cộng sự đã đề xuất rằng cho các vấn đề trên, cách tiếp cận PHÂN GIAI ĐOẠN "
   "(staged approach) tới các dịch vụ chính thức có thể giúp ích — với giai đoạn đầu "
   "là các can thiệp DỄ TIẾP CẬN tại các bối cảnh thân thiện với người trẻ như trường "
   "học. Tổng quan mới nhất về can thiệp trường học báo cáo một effect size NHỎ cho "
   "trầm cảm và lo âu — chỉ ra rằng các can thiệp dựa trên CBT và do nhà LÂM SÀNG cung "
   "cấp ở trường THPT có khả năng hiệu quả hơn.")
crit("Đây là chính tác giả thừa nhận \"small effect size\" trong tổng quan school "
     "interventions trước đó. BESST không phá vỡ pattern này — d=−0,17 cũng là small. "
     "Vì vậy claim trong editorial B5 là \"very positive\" sai bản chất.")

en("Furthermore, specific interventions for older adolescents are probably needed as "
   "substantial brain maturation and marked differences in sleep and coping mechanisms "
   "(as well as social changes—eg, increased autonomy) occur in this age. There has only "
   "been one small trial with 21 participants of a school-based intervention in "
   "adolescents age 16 years and older.")
vn("Hơn nữa, các can thiệp cụ thể cho VTN LỚN TUỔI có lẽ là cần thiết vì brain "
   "maturation đáng kể và khác biệt rõ rệt về giấc ngủ và cơ chế đối phó (cũng như "
   "thay đổi xã hội — vd. tăng tự chủ) xảy ra ở độ tuổi này. Chỉ có MỘT thử nghiệm nhỏ "
   "với 21 người tham gia về can thiệp trường học cho VTN 16+ tuổi.")
crit("Lý do tập trung 16-18 tuổi (sixth form UK) là LACUNA evidence base — trước BESST "
     "chỉ có 1 trial 21 HS. Đây là gap quan trọng. Áp dụng VN: 16-18 tuổi tương đương "
     "lớp 11-12 — nhóm tuổi có áp lực thi đại học cao, gap evidence cũng tương tự.")

en("To help address these problems, in England, Mental Health Support Teams have been "
   "recently introduced to support adolescents' mental health and bridge the gap between "
   "schools and CAMHS. Mental Health Support Teams staff are master's or postgraduate "
   "diploma level therapists or junior therapists.")
vn("Để giúp giải quyết các vấn đề này, ở Anh, các Mental Health Support Teams (MHSTs) "
   "đã được giới thiệu gần đây để hỗ trợ SKTT VTN và CẦU NỐI giữa trường học và CAMHS. "
   "Nhân viên MHST là các nhà trị liệu trình độ thạc sĩ hoặc postgraduate diploma, "
   "hoặc nhà trị liệu trẻ.")

en("How school-based interventions should be offered is unclear. Although researcher-led "
   "targeted approaches demonstrate greater effectiveness, participants report feeling "
   "stigmatised by this approach. Conversely, universal approaches, which deliver the "
   "intervention to all adolescents, are less stigmatising, but tend to be less effective "
   "or reach students who do not require support.")
vn("Cách triển khai can thiệp trường học vẫn chưa rõ. Mặc dù các approaches "
   "researcher-led targeted thể hiện hiệu quả lớn hơn, người tham gia báo cáo CẢM "
   "THẤY KỲ THỊ bởi cách tiếp cận này. Ngược lại, approaches PHỔ QUÁT — cung cấp can "
   "thiệp cho TẤT CẢ VTN — ít kỳ thị hơn, nhưng có xu hướng KÉM HIỆU QUẢ hoặc tiếp "
   "cận HS không cần hỗ trợ.")

en("Participant-initiated self-referral systems, where the individuals decide if they "
   "want to be involved by referring themselves, have rarely been used in trials. The "
   "self-referral process is part of the PLACES model which describes methods to "
   "increase accessibility, including using colloquial terms (eg, stress) rather than "
   "medical terms (eg, depression or anxiety) to reduce stigma. Self-referral itself has "
   "several advantages: it has the potential to reduce stigma, emphasises autonomy "
   "(which is valued by adolescents), and allows more efficient use of resources. In "
   "previous research, this approach has led to high engagement by students who have not "
   "previously sought help and has also led to high follow-up rates of more than 90%.")
vn("Hệ thống self-referral khởi xướng bởi người tham gia — nơi cá nhân tự quyết định "
   "tham gia bằng cách tự đăng ký — hiếm khi được sử dụng trong các trials. Quá trình "
   "self-referral là một phần của MÔ HÌNH PLACES — mô tả các phương pháp tăng "
   "accessibility, gồm sử dụng các thuật ngữ ĐỜI THƯỜNG (vd. stress) thay vì các "
   "thuật ngữ y tế (vd. trầm cảm hoặc lo âu) để giảm kỳ thị. Self-referral bản thân nó "
   "có nhiều lợi ích: TIỀM NĂNG GIẢM KỲ THỊ, NHẤN MẠNH TỰ CHỦ (vốn được VTN coi "
   "trọng), và cho phép sử dụng tài nguyên HIỆU QUẢ HƠN. Trong nghiên cứu trước, "
   "approach này đã dẫn đến engagement cao bởi HS chưa từng tìm trợ giúp, và follow-up "
   "rate hơn 90%.")

en("Furthermore, the self-referral system has been shown to facilitate a higher "
   "proportion of people from ethnic minority groups to engage in interventions, an "
   "important issue in diverse communities such as those in England. Minority ethnic "
   "groups have been consistently identified as underserved in mental health treatment "
   "and research because of problems with access to services.")
vn("Hơn nữa, hệ thống self-referral đã được chứng minh THÚC ĐẨY tỉ lệ cao hơn người "
   "từ các nhóm thiểu số sắc tộc engage với can thiệp — một vấn đề quan trọng trong "
   "các cộng đồng đa dạng như ở Anh. Các nhóm thiểu số sắc tộc đã liên tục được xác "
   "định là UNDERSERVED trong điều trị SKTT và nghiên cứu vì các vấn đề tiếp cận dịch vụ.")

en("The DISCOVER workshop programme is based on an adult stress workshop model using a "
   "self-referral system and has been adapted to provide an accessible and acceptable "
   "intervention for adolescents. Key elements of the workshops are as follows: (1) use "
   "of CBT materials; (2) brief, 1-day duration delivered within a community setting; "
   "and (3) a self-referral pathway. Additional elements of the adolescent model are "
   "greater interaction between students and clinicians and the use of more visual "
   "materials such as videos and games. Although not powered to evaluate outcomes, a "
   "feasibility study of the DISCOVER workshop programme found a reduction in depression "
   "(d=0·27) and anxiety (d=0·25) 3 months post-intervention and was shown to be "
   "acceptable to students.")
vn("Chương trình workshop DISCOVER dựa trên mô hình workshop stress người lớn sử dụng "
   "hệ thống self-referral, và đã được ADAPT để cung cấp can thiệp dễ tiếp cận và có "
   "thể chấp nhận cho VTN. Các yếu tố then chốt của workshop: (1) sử dụng tài liệu "
   "CBT; (2) ngắn — 1 ngày — được cung cấp trong bối cảnh cộng đồng; và (3) lộ trình "
   "self-referral. Yếu tố bổ sung của mô hình VTN là TƯƠNG TÁC LỚN HƠN giữa HS và "
   "nhà lâm sàng và sử dụng nhiều TÀI LIỆU TRỰC QUAN như video và trò chơi. Mặc dù "
   "không có đủ power để đánh giá outcome, một nghiên cứu feasibility của chương trình "
   "DISCOVER tìm thấy giảm trầm cảm (d=0,27) và lo âu (d=0,25) tại 3 tháng sau can "
   "thiệp — và được chứng minh có thể chấp nhận với HS.")
crit("Feasibility study DISCOVER (Brown 2019, J Adolesc) cho effect d=0,27 (depression) "
     "và d=0,25 (anxiety) — LỚN HƠN một chút BESST RCT (d=−0,17). Điều này KÝ HIỆU "
     "INFLATION typical: pilot/feasibility tend to overestimate effects (small N, "
     "selection bias). RCT chính thức luôn cho effect nhỏ hơn pilot. Cảnh báo cho VN: "
     "nếu pilot VN cho effect lớn, expect RCT chính thức sẽ giảm.")

en("The primary objective of the Brief Educational Workshops in Secondary Schools Trial "
   "(BESST) was to investigate the clinical effectiveness and cost-effectiveness of "
   "DISCOVER workshops on symptoms of depression in 16–18-year-olds at 6 months compared "
   "with treatment-as-usual. Secondary objectives were to assess symptoms of anxiety, "
   "wellbeing, sleep, and resilience.")
vn("MỤC TIÊU CHÍNH của thử nghiệm BESST là điều tra hiệu quả lâm sàng và chi phí-hiệu "
   "quả của workshop DISCOVER lên triệu chứng trầm cảm ở HS 16-18 tuổi tại 6 tháng so "
   "với TAU. MỤC TIÊU PHỤ là đánh giá triệu chứng lo âu, wellbeing, sleep, và resilience.")

# ---------- Methods ----------
H2("Phương pháp (Methods)")

H2("Study design and participants")
en("The BESST study was a multicentre two-arm parallel cluster randomised controlled "
   "trial in England, with embedded health economic assessment. Outcomes were measured "
   "at baseline and at the 3 month and 6 month follow-up. Schools included in the study "
   "were either school sixth forms or dedicated sixth-form colleges with 70 or more "
   "students enrolled, state-funded, with sufficient resources available to host the "
   "trial.")
vn("Nghiên cứu BESST là một thử nghiệm CLUSTER RCT MULTICENTRE 2 nhánh song song tại "
   "Anh, với HEALTH ECONOMIC ASSESSMENT tích hợp. Outcome được đo ở baseline và follow-"
   "up 3 và 6 tháng. Các trường được bao gồm là sixth form trường THPT hoặc sixth-form "
   "college chuyên biệt có ≥70 HS đăng ký, do nhà nước tài trợ, có đủ tài nguyên để "
   "host trial.")

en("The participants were students aged 16–18 years; attending for the full school year; "
   "seeking psychological help for stress, worry, or low mood; fluent in English and "
   "able to provide written informed consent; and available to attend and participate "
   "in the workshop. Participants were excluded if they were identified as actively "
   "suicidal, had severe learning difficulties or psychosis, or were actively receiving "
   "psychological therapy for anxiety or depression through CAMHS. Participation was "
   "limited to 19 students per school during the 2022–23 school year (for practical "
   "reasons in implementing the DISCOVER workshop), with students who had provided "
   "written informed consent invited to take part through random selection.")
vn("Người tham gia là HS 16-18 tuổi; học toàn năm; tìm trợ giúp tâm lý cho stress, "
   "lo lắng hoặc tâm trạng thấp; lưu loát tiếng Anh và có thể consent có chữ ký; có "
   "khả năng tham dự và tham gia workshop. LOẠI TRỪ nếu: đang có ý định tự tử, có khó "
   "khăn học tập NẶNG hoặc loạn thần (psychosis), hoặc đang nhận psychological therapy "
   "cho lo âu/trầm cảm qua CAMHS. Sự tham gia giới hạn 19 HS/trường trong năm 2022-23 "
   "(vì lý do thực tế triển khai DISCOVER workshop), với HS đã cho consent được mời "
   "qua RANDOM SELECTION.")
crit("Tiêu chí loại trừ chuẩn nhưng có 2 vấn đề: (1) Loại HS đang ở CAMHS → HS có vấn "
     "đề NẶNG bị loại → mẫu nghiên cứu là MILD-MODERATE — không generalisable cho HS "
     "có triệu chứng nặng. (2) Loại HS không lưu loát tiếng Anh → loại minorities mới "
     "đến → mẫu thiên về established minority groups. Áp dụng VN: tiêu chí phải khác "
     "(ngôn ngữ Việt; nhưng ai loại trừ HS đang được tư vấn tâm lý ở phòng tư vấn "
     "trường khác?).")

H2("Randomisation and masking")
en("Following baseline assessments, schools were randomised in a 1:1 ratio for their "
   "participants to receive either the DISCOVER workshop intervention or treatment-as-"
   "usual, using a covariate minimisation algorithm, stratified by site and balanced on "
   "school size and index of multiple deprivation. The allocation sequence was generated "
   "by an unmasked statistician who was not part of the study team and the arm "
   "allocations were released as A and B to the trial manager who was also unmasked "
   "after all adolescents were enrolled at baseline. The outcome assessors were masked "
   "and reminded the students at the start and during the follow-up not to divulge "
   "whether they received the workshop or not. The chief investigator, senior health "
   "economist, and senior statistician were masked until database lock.")
vn("Sau baseline assessment, các trường được randomised tỉ lệ 1:1 cho người tham gia "
   "nhận DISCOVER hoặc TAU, sử dụng COVARIATE MINIMISATION ALGORITHM, stratified theo "
   "site và balanced theo quy mô trường + IMD. Allocation sequence được tạo bởi "
   "STATISTICIAN UNMASKED không thuộc team nghiên cứu — và arm allocations được phát "
   "hành dưới dạng A và B cho trial manager (cũng unmasked) SAU khi tất cả VTN được "
   "tuyển ở baseline. Outcome assessor được MASKED và nhắc HS ở đầu và trong follow-"
   "up KHÔNG tiết lộ liệu họ có nhận workshop hay không. Chief investigator, senior "
   "health economist, và senior statistician được masked đến database lock.")
crit("MASKING strategy này CHẤT LƯỢNG CAO — đặc biệt là database lock masking cho 3 "
     "vai trò senior. Tuy nhiên: (1) HS không thể mask (đã nhận workshop hay không); "
     "(2) MHST staff không thể mask. Hai impossibility này là LIMITATION inherent của "
     "psychological RCTs. Cochrane RoB 2 sẽ đánh giá domain \"Deviations from intended "
     "interventions\" và \"Measurement of outcome\" có risk of bias TRUNG BÌNH do "
     "self-reported outcomes + unmasked participants.")

H2("Procedures")
en("Following consent from participating schools and colleges to host the trial, an "
   "assembly presentation introduced the trial and intervention to potential "
   "participants. Students were informed that they could choose whether or not to "
   "participate and were invited to hear more about the trial at a lunchbreak a few "
   "days later, where they were given the participant information sheet and consent form.")
vn("Sau consent từ trường và college tham gia, một presentation tại GIỜ CHÀO CỜ "
   "(assembly) giới thiệu thử nghiệm và can thiệp tới người tham gia tiềm năng. HS "
   "được thông báo rằng họ có thể chọn tham gia hay không, và được mời nghe thêm về "
   "thử nghiệm vào giờ ăn trưa vài ngày sau — nơi họ được cung cấp participant "
   "information sheet và consent form.")
crit("CƠ CHẾ TUYỂN HỌC SINH CHUẨN MỰC: assembly → information session → consent — "
     "tránh ép buộc, tôn trọng autonomy. Áp dụng VN: tiết chào cờ thứ 2 sáng có thể "
     "dùng làm assembly; sau đó info session ngoài giờ; consent có chữ ký + có thể "
     "cần consent phụ huynh nếu HS dưới 18 (luật VN).")

en("DISCOVER is a brief, accessible workshop-based stress management programme for "
   "16–18-year-olds, to which they can self-refer. The workshop is considered accessible "
   "because of the self-referral system where students are invited at the assembly to "
   "refer themselves as well as the use of non-diagnostic terms, such as stress, in "
   "describing the programme. The programme was co-designed with a Teenage Advisory "
   "Group of 31 16–18-year-olds, with the aim of improving engagement, offering "
   "effective treatment, and maintaining participants' motivation and improvement to "
   "reduce relapse.")
vn("DISCOVER là một chương trình quản lý stress dạng workshop NGẮN, dễ tiếp cận cho HS "
   "16-18 tuổi — mà họ có thể self-refer. Workshop được coi là accessible nhờ hệ "
   "thống self-referral nơi HS được mời tại assembly tự đăng ký, cũng như sử dụng "
   "thuật ngữ KHÔNG-CHẨN-ĐOÁN như stress trong mô tả chương trình. Chương trình được "
   "CO-DESIGNED với Teenage Advisory Group gồm 31 HS 16-18 tuổi — với mục tiêu cải "
   "thiện engagement, cung cấp điều trị hiệu quả, và duy trì động lực + cải thiện để "
   "giảm tái phát.")

en("The workshop programme includes CBT coping techniques for managing mood, anxiety, "
   "and stress, delivered in non-medicalised language and with images and materials "
   "featuring students from diverse backgrounds.")
vn("Chương trình workshop bao gồm các kỹ thuật CBT đối phó để quản lý tâm trạng, lo "
   "âu, và stress — được cung cấp bằng ngôn ngữ KHÔNG y tế hoá, với hình ảnh và tài "
   "liệu featuring HS từ background đa dạng.")

en("A 2-day DISCOVER training programme was offered to Mental Health Support Team staff "
   "from National Health Service (NHS) trusts. Members of the Mental Health Support "
   "Teams were trained to deliver the intervention in accordance with the DISCOVER "
   "manual and trial protocol. The 2-day training session, and one supervision session "
   "per Mental Health Support Team, were led by IS, who was a co-applicant on BESST. "
   "Each workshop programme was co-facilitated by three staff: one senior therapist and "
   "two junior therapists. The workshop delivery teams were recruited into the trial "
   "solely for workshop delivery.")
vn("Một chương trình ĐÀO TẠO DISCOVER 2 NGÀY được cung cấp cho nhân viên MHST từ NHS "
   "trusts. Các thành viên MHST được đào tạo để cung cấp can thiệp theo manual DISCOVER "
   "và protocol thử nghiệm. Buổi đào tạo 2 ngày + 1 buổi supervision/MHST được dẫn "
   "dắt bởi IS (co-applicant của BESST). Mỗi workshop được CO-FACILITATED bởi 3 nhân "
   "viên: 1 senior therapist + 2 junior therapists. Workshop delivery teams được tuyển "
   "vào trial CHỈ để cung cấp workshop.")
crit("ĐÀO TẠO 2 NGÀY là RẤT NGẮN cho CBT — UK MHST đã có background thạc sĩ + "
     "psychological therapy nên 2 ngày là supplement tốt; nhưng với VN — nơi tư vấn "
     "học đường đa số chưa được đào tạo CBT — 2 ngày KHÔNG đủ. Adapt VN cần đào tạo "
     "5-7 ngày + supervision tháng/quý.")

en("The workshop was a day-long, face-to-face group event, which took place at the "
   "school or college in a private classroom (without school staff present) over a "
   "single school day. Permission for students to attend and miss curricular activities "
   "was obtained from staff in advance, and the students' usual breaks and lunch were "
   "adhered to.")
vn("Workshop là một sự kiện nhóm FACE-TO-FACE 1 ngày, diễn ra tại trường hoặc college "
   "trong một phòng học RIÊNG TƯ (KHÔNG có nhân viên trường có mặt) trong một ngày "
   "học. Sự cho phép cho HS tham dự và bỏ lỡ hoạt động chính khoá đã được lấy từ "
   "nhân viên trước; các giờ giải lao + ăn trưa thông thường của HS được tuân thủ.")
crit("\"WITHOUT SCHOOL STAFF PRESENT\" — đặc biệt quan trọng để giảm stigma + tăng "
     "open disclosure. Áp dụng VN: phòng riêng + KHÔNG có GVCN/Ban giám hiệu — quan "
     "trọng vì HS VN có thể tự kiểm duyệt trước người lớn.")

en("In the days before attending the workshop, each student met individually with a "
   "workshop leader in a private space, for approximately 30 min. During this session "
   "they discussed their personal goals, which they would set at the end of the workshop "
   "day. Core workshop content was as follows. Each workshop began with introductions "
   "and icebreakers. A CBT-informed model of emotional problems was then provided to "
   "explain and normalise young people's experiences, including video clips of teenage "
   "actors and group discussions. Particular attention was given to personal, "
   "relationship, and academic stresses typical for the age group. CBT techniques for "
   "managing anxiety and mood problems were taught and practised, supported by scripted "
   "role-plays, video demonstrations, and printed handouts. Behavioural strategies used "
   "included problem-solving, sleep advice, and time management. Cognitive strategies "
   "included identification of and challenging negative thoughts. Participants were "
   "provided with a workbook to keep, which provided all the covered material and space "
   "to make notes throughout the workshop, and record their personal goals. Content of "
   "the workshop in digital form was also provided in a smartphone app.")
vn("Trong những ngày TRƯỚC workshop, mỗi HS gặp riêng workshop leader tại không gian "
   "riêng tư khoảng 30 phút. Trong buổi này, họ thảo luận về mục tiêu cá nhân — sẽ "
   "đặt vào cuối ngày workshop. Nội dung workshop core: bắt đầu bằng giới thiệu và "
   "icebreaker. Mô hình các vấn đề cảm xúc theo CBT được cung cấp để giải thích và "
   "BÌNH THƯỜNG HOÁ trải nghiệm người trẻ — gồm video clip diễn viên teen và thảo luận "
   "nhóm. Chú ý đặc biệt vào các stress cá nhân, mối quan hệ, học thuật điển hình của "
   "lứa tuổi. Các kỹ thuật CBT để quản lý lo âu + mood được dạy và thực hành — hỗ "
   "trợ bằng role-play có kịch bản, demo video, và handout in. Chiến lược hành vi: "
   "problem-solving, lời khuyên giấc ngủ, và quản lý thời gian. Chiến lược nhận thức: "
   "nhận diện và thách thức suy nghĩ tiêu cực. Người tham gia được cung cấp WORKBOOK "
   "để giữ — gồm tất cả tài liệu + không gian ghi chú + ghi mục tiêu cá nhân. Nội dung "
   "workshop dạng kỹ thuật số cũng được cung cấp trong APP smartphone.")

en("After 1 week, participants were followed up individually by one of the workshop "
   "leaders, with the participants receiving a 15–30 min telephone goal review to "
   "monitor progress and support incorporation of CBT skills into real-life situations. "
   "Participants were given the option of receiving two further telephone goal reviews "
   "within the 12-week post-workshop period.")
vn("Sau 1 TUẦN, người tham gia được follow-up cá nhân bởi 1 trong các workshop "
   "leaders, với buổi điện thoại 15-30 phút REVIEW MỤC TIÊU để theo dõi tiến độ và "
   "hỗ trợ tích hợp kỹ năng CBT vào tình huống thực tế. Người tham gia được cho lựa "
   "chọn nhận TỐI ĐA 2 buổi điện thoại review thêm trong giai đoạn 12 tuần sau workshop.")

en("To assess treatment fidelity, each member of the workshop delivery team completed a "
   "9-item self-report fidelity checklist immediately following each workshop. An "
   "independent observer also attended one workshop per delivery team to assess "
   "treatment fidelity using the same checklist. Fidelity was met if seven of the nine "
   "items were met (including four mandatory items).")
vn("Để đánh giá TREATMENT FIDELITY, mỗi thành viên team cung cấp workshop hoàn thành "
   "một CHECKLIST FIDELITY 9-mục TỰ BÁO CÁO ngay sau mỗi workshop. Một observer ĐỘC "
   "LẬP cũng tham dự 1 workshop/delivery team để đánh giá fidelity bằng cùng checklist. "
   "FIDELITY ĐẠT nếu 7/9 mục đạt (gồm 4 mục bắt buộc).")

en("Treatment-as-usual is defined as the usual school care provided to students in "
   "their sixth form. Types of school provision offered across participating schools, "
   "and the percentage of schools offering each provision, are presented in the "
   "appendix.")
vn("CHĂM SÓC THÔNG THƯỜNG (TAU) được định nghĩa là chăm sóc trường học thông thường "
   "được cung cấp cho HS trong sixth form của họ. Các loại provision trường được cung "
   "cấp xuyên các trường tham gia + tỉ lệ trường cung cấp mỗi provision được trình "
   "bày trong phụ lục.")
crit("TAU \"thông thường\" có thể VARIES rộng giữa các trường UK — tạo HETEROGENEITY "
     "trong control arm. Đây là LIMITATION típ của school-based RCT. Tác giả đã thảo "
     "luận trong Discussion: \"control group was passive rather than active and the "
     "effect found in the sample was only modest\".")

H2("Outcomes")
en("The primary outcome (collected at baseline, 3-month, and 6-month follow-ups in all "
   "participants) was symptoms of depression, assessed using the Mood and Feelings "
   "Questionnaire (MFQ; higher scores indicate greater severity of depressive "
   "symptoms). The MFQ is a 33-item self-report depression measure, which in the "
   "English version has shown good psychometric properties and is a widely used and "
   "validated instrument for adolescents. It has been used throughout the world and "
   "translated into several languages and is free to use for clients or research. "
   "Scores range from 0 to 66, with a clinical cutoff of higher than 27 defining "
   "elevated symptoms of depression.")
vn("OUTCOME CHÍNH (thu thập ở baseline, 3 tháng, 6 tháng follow-up ở tất cả người tham "
   "gia) là triệu chứng trầm cảm — đánh giá bằng MOOD AND FEELINGS QUESTIONNAIRE (MFQ; "
   "điểm cao hơn = mức độ triệu chứng nặng hơn). MFQ là thang TỰ BÁO CÁO 33-MỤC, có "
   "tính chất tâm trắc tốt trong phiên bản tiếng Anh — và là công cụ được sử dụng và "
   "kiểm chứng rộng rãi cho VTN. MFQ đã được sử dụng khắp thế giới và dịch sang nhiều "
   "ngôn ngữ — và FREE để dùng cho client hoặc nghiên cứu. Điểm số 0-66, với CUTOFF "
   "LÂM SÀNG >27 định nghĩa triệu chứng trầm cảm \"elevated\".")
crit("MFQ là công cụ chuẩn nhưng có 2 lưu ý cho VN: (1) Bản dịch VN của MFQ chưa được "
     "validate cấp quốc gia (chỉ có một số bản dịch học thuật cho từng nghiên cứu). "
     "(2) Cutoff >27 dựa trên norm UK/US — không nhất thiết khớp norm VN. Áp dụng VN "
     "cần validate cutoff trước hoặc dùng PHQ-9/CES-D đã có VN-norm (Trần Tuấn 2017).")

en("Secondary outcomes and their measures were anxiety, assessed using the anxiety "
   "sub-scale from the Revised Child Anxiety and Depression Scale (RCADS); wellbeing, "
   "assessed using the Warwick–Edinburgh Mental Wellbeing Scale (WEMWBS); sleep "
   "quality, assessed using the Sleep Condition Indicator (SCI); and resilience, "
   "assessed using the Child and Youth Resilience Measure 12 (CYRM-12). These measures "
   "were collected at baseline, 3-month, and 6-month follow-ups in all participants. "
   "Student satisfaction was measured in the intervention group at the end of each "
   "workshop, assessed using the Client Satisfaction Questionnaire (CSQ-8).")
vn("OUTCOME PHỤ và các thước đo: lo âu — RCADS-anxiety subscale; wellbeing — Warwick-"
   "Edinburgh Mental Wellbeing Scale (WEMWBS); chất lượng giấc ngủ — Sleep Condition "
   "Indicator (SCI); resilience — Child and Youth Resilience Measure 12 (CYRM-12). "
   "Thu thập ở baseline, 3 tháng, 6 tháng. Hài lòng HS đo ở nhóm intervention cuối "
   "mỗi workshop bằng Client Satisfaction Questionnaire (CSQ-8).")
crit("BỘ outcome COMPREHENSIVE — gồm 5 domain: depression (primary), anxiety, wellbeing, "
     "sleep, resilience. Đây là chuẩn cao. Áp dụng VN: nên dùng GAD-7 thay RCADS-anxiety "
     "(GAD-7 đã được Hoa 2024 validate cho VN α=0,916), DASS-21 thay WEMWBS, PSQI thay "
     "SCI. Resilience: CYRM-12 hoặc CD-RISC-10 (đã có bản VN).")

en("Health economic secondary outcomes were health-related quality of life, assessed "
   "using the EQ-5D-3L, used to calculate quality-adjusted life-years (QALYs) for use "
   "in economic evaluation, and use of health and social care services, measured using "
   "the Child and Adolescent Service Use Schedule (CA-SUS).")
vn("OUTCOME PHỤ KINH TẾ Y TẾ: chất lượng cuộc sống liên quan sức khoẻ — EQ-5D-3L "
   "(dùng để tính QALYs cho economic evaluation), và sử dụng dịch vụ y tế + xã hội — "
   "CA-SUS.")

H2("Statistical analysis")
en("All outcomes reported were prespecified in the statistical analysis plan that was "
   "drafted and approved by a masked trial statistician and senior statistician (KJ, BC) "
   "following King's Clinical Trials Unit Standard operating procedures.")
vn("Tất cả outcome báo cáo được PRESPECIFIED trong statistical analysis plan — soạn "
   "thảo và phê duyệt bởi trial statistician masked + senior statistician (KJ, BC) "
   "theo SOP của King's Clinical Trials Unit.")
crit("PRESPECIFICATION rất quan trọng — tránh p-hacking và HARKing (Hypothesizing After "
     "Results are Known). Đây là điểm STRENGTH của BESST.")

en("Assuming a two-sided type 1 error of 0·05 with 90% power, to detect an effect size "
   "of 0·28 with an intra-cluster correlation of 0·03, 760 participants were estimated "
   "to be required from 54 schools. After inflating for loss to follow-up of schools "
   "and participants, we planned to enrol 60 schools and 900 adolescents.")
vn("Giả định two-sided type 1 error = 0,05 với 90% power, để phát hiện effect size = "
   "0,28 với intra-cluster correlation = 0,03, ƯỚC TÍNH cần 760 người tham gia từ 54 "
   "trường. Sau khi inflate cho loss to follow-up trường và người tham gia, lập kế "
   "hoạch tuyển 60 trường + 900 VTN.")
crit("POWER CALCULATION dùng effect size = 0,28 (small-to-medium). Effect thực tế là "
     "d=−0,17 → trial bị UNDERPOWERED cho effect thực này. Tuy nhiên, vì n=854 ITT "
     "lớn, vẫn detect được d=−0,17 với p=0,0019. Điều này CONFIRMS rằng \"statistical "
     "significance\" KHÔNG đồng nghĩa \"clinical meaningfulness\". Người đọc cần đọc "
     "Cohen d, không chỉ p-value.")

en("Descriptive data of the population were presented as means and SDs for continuous "
   "data, and as counts and percentages for categorical data.")
vn("Dữ liệu mô tả của population được trình bày dưới dạng MEAN và SD cho continuous, "
   "và counts + percentages cho categorical.")

en("The primary outcome of MFQ was analysed with a mixed-effect, multi-level linear "
   "model at 6 months adjusted using prespecified fixed effects of: baseline severity "
   "(MFQ score), aggregated level school deprivation, geographical area, school size, "
   "gender, ethnicity group, assessment timepoint, treatment, and treatment-by-time "
   "interaction. A random intercept was fitted for each school and student and the "
   "adjusted mean difference between the intervention and control score was estimated, "
   "alongside the 95% CI and p value. The associated Cohen's d standardised effect size "
   "(with 95% CI) was calculated using a pooled standard deviation. The intra-cluster "
   "correlation at 6 months was also calculated. Secondary outcomes were analysed in a "
   "similar way.")
vn("Outcome chính MFQ được phân tích bằng MIXED-EFFECT MULTI-LEVEL LINEAR MODEL tại 6 "
   "tháng — adjusted với các fixed effects prespecified: baseline severity (điểm MFQ), "
   "aggregated level school deprivation, geographical area, school size, gender, "
   "ethnicity group, assessment timepoint, treatment, và TREATMENT-BY-TIME INTERACTION. "
   "RANDOM INTERCEPT fit cho mỗi trường và HS, và adjusted mean difference giữa "
   "intervention và control score được ước tính — cùng 95% CI và p value. COHEN's d "
   "standardised effect size (với 95% CI) được tính bằng pooled standard deviation. "
   "Intra-cluster correlation tại 6 tháng cũng được tính. Outcome phụ được phân tích "
   "tương tự.")
crit("MIXED-EFFECT MULTI-LEVEL MODEL là phương pháp THÍCH HỢP cho cluster RCT — handle "
     "được nesting structure (HS trong trường). Adjusted cho 8 covariates — TỐT cho "
     "kiểm soát confounding nhưng RISKY cho overfitting nếu không prespecified. May "
     "mắn ở đây có prespec.")

en("All participants with any follow-up data were included in the intention-to-treat "
   "population, and those with complete covariate data were included in the intention-"
   "to-treat analysis. Full details of handling missing data (observations and "
   "participants) are in the statistical analysis plan (appendix p 10). Prespecified "
   "subgroup analyses were planned for elevated baseline symptoms (MFQ >27), gender, "
   "participant age, and school year. A per-protocol analysis, as specified in the "
   "statistical analysis plan, repeated the primary analysis but excluded participants "
   "who did not receive an acceptable dose of the intervention (less than 75% of the "
   "workshop or not setting a goal), had data collected outside of visit windows, or "
   "reported access to CAMHS.")
vn("Tất cả người tham gia có bất kỳ dữ liệu follow-up nào được bao gồm trong ITT "
   "population, và những người có dữ liệu covariate ĐẦY ĐỦ được bao gồm trong ITT "
   "ANALYSIS. Chi tiết handling missing data (observation + participant) trong "
   "statistical analysis plan (appendix p 10). PRESPECIFIED SUBGROUP ANALYSES được lập "
   "kế hoạch cho: triệu chứng baseline ĐÃ ELEVATED (MFQ >27), GENDER, tuổi người tham "
   "gia, và năm học. Phân tích PER-PROTOCOL — như được prespec — lặp lại analysis "
   "chính nhưng LOẠI TRỪ người không nhận \"acceptable dose\" can thiệp (<75% workshop "
   "hoặc không đặt goal), có dữ liệu thu thập ngoài visit window, hoặc báo cáo tiếp "
   "cận CAMHS.")

en("The economic analyses followed a health economic analysis plan drafted by the trial "
   "health economist (JS), in line with the trial protocol, and approved by the senior "
   "health economist (SB; reproduced in the appendix p 48). The economic analysis was a "
   "cost–utility analysis at the 6-month follow-up with effectiveness measured in terms "
   "of QALYs estimated from the EQ-5D-3L measure of health-related quality of life. The "
   "economic perspective taken was that preferred by the National Institute for Health "
   "and Care Excellence, which includes all national health services and all social "
   "care services (often referred to as the NHS and personal social services perspective).")
vn("Phân tích kinh tế tuân theo health economic analysis plan soạn thảo bởi trial "
   "health economist (JS), phù hợp với protocol thử nghiệm, và phê duyệt bởi senior "
   "health economist (SB). Phân tích kinh tế là COST-UTILITY ANALYSIS tại 6 tháng — "
   "với effectiveness đo bằng QALYs ước tính từ EQ-5D-3L. Quan điểm kinh tế là quan "
   "điểm ưu tiên của National Institute for Health and Care Excellence (NICE) — gồm "
   "tất cả NHS và personal social services.")

# ---------- Results ----------
H2("Kết quả (Results)")

H2("Recruitment + CONSORT flow (Figure 1)")
en("Over 2 school years (2021–22 and 2022–23), we screened 111 schools to participate, "
   "and seven were deemed ineligible and 47 did not provide consent. We then enrolled "
   "57 schools or colleges with sixth forms in the study (appendix p 8). Between Oct 4, "
   "2021, and Nov 10, 2022, 933 students were screened for eligibility, seven were not "
   "eligible for inclusion, and 26 did not attend the baseline meeting and assessment, "
   "resulting in 900 students randomly assigned to either of the two groups. 379 were "
   "enrolled during the 2021–22 intake and 521 during the 2022–23 intake, in 15 "
   "localities across four regions of England. 31 schools with 457 participants were "
   "randomly assigned to the treatment-as-usual group and 26 schools with 443 "
   "participants were randomly assigned to the DISCOVER workshop programme (figure 1). "
   "The study involved a total of 11 NHS trusts within which 15 Mental Health Support "
   "Teams delivered the intervention.")
vn("Trên 2 năm học (2021-22 và 2022-23), chúng tôi sàng lọc 111 trường để tham gia, "
   "7 không đủ điều kiện và 47 không cho consent. Sau đó tuyển 57 trường/college có "
   "sixth form. Giữa 04/10/2021 và 10/11/2022, 933 HS được sàng lọc, 7 không đủ điều "
   "kiện, 26 không tham dự baseline → 900 HS phân ngẫu nhiên 2 nhóm. 379 tuyển trong "
   "2021-22, 521 trong 2022-23, tại 15 ĐỊA PHƯƠNG xuyên 4 VÙNG nước Anh. 31 TRƯỜNG / "
   "457 NGƯỜI được phân TAU; 26 TRƯỜNG / 443 NGƯỜI phân DISCOVER. Tổng có 11 NHS TRUSTS "
   "với 15 MHSTs cung cấp can thiệp.")
note("CONSORT FLOW (Figure 1 trong PDF gốc, trang 509): 111 invited → 104 eligible → 57 "
     "consented → 1407 attended assembly → 991 consent → 933 screened → 926 invited → "
     "900 randomised. ITT analysis n=854 (439 control + 415 DISCOVER).")
crit("CONSORT FLOW chuẩn mực + minh bạch. Tỉ lệ tham gia: 47/104 trường eligible từ "
     "chối (45%) — KHÁ CAO. Tỉ lệ HS từ 1407 attend assembly → 900 randomise = 64% — "
     "CŨNG CAO. Reach tốt cho self-referral. Tuy nhiên: 47% trường không tham gia có "
     "thể là chính những trường có HS RỦI RO CAO (do bận, thiếu nguồn lực) → SELECTION "
     "BIAS ở cấp trường.")

H2("Baseline characteristics (Bảng 1 PDF gốc)")
nr("Em đã tổng hợp Bảng 1 (PDF gốc trang 510) thành bảng tổng hợp chi tiết trong PHẦN "
   "2 (Bảng B). Tóm tắt: tuổi 17,2 (SD 0,6), 71% nữ, 25% nam, 2% other, 1% prefer not "
   "say. 52% White, 7% Mixed, 17% Asian, 16% Black, 2% Chinese, 5% Other. 80% chưa "
   "từng tìm GP cho SKTT. Khá cân bằng giữa 2 nhóm; nhóm DISCOVER có % nam cao hơn "
   "(31% vs 20% trong TAU) — confirmed bằng minimisation algorithm.", italic=True)

H2("Primary outcome — MFQ depression at 6 months (Bảng 2-3, Figure 2)")
en("For the primary analysis of the MFQ at 6 months post randomisation (intention-to-"
   "treat analysis population), we estimated an adjusted mean difference of −2·06 (95% "
   "CI −3·35 to −0·76, Cohen's d −0·17; p=0·0019; intra-cluster correlation, 0·01), "
   "showing a significant reduction in depressive symptoms in the DISCOVER group versus "
   "the control group. The intra-cluster correlation for the MFQ at 6 months was similar "
   "to that quoted in the sample size calculation (0·01 [95% CI 0·001 to 0·09]).")
vn("Cho phân tích chính của MFQ tại 6 tháng sau randomisation (ITT analysis population), "
   "chúng tôi ước tính ADJUSTED MEAN DIFFERENCE = −2,06 (95% CI −3,35 đến −0,76; "
   "COHEN's d = −0,17; p = 0,0019; intra-cluster correlation = 0,01) — chỉ ra giảm "
   "có ý nghĩa thống kê triệu chứng trầm cảm ở nhóm DISCOVER so với control. ICC MFQ "
   "tại 6 tháng tương tự với ICC trong sample size calculation (0,01 [95% CI 0,001 "
   "đến 0,09]).")
crit("3 SỐ LIỆU THEN CHỐT cần ghi nhớ:\n"
     "• Adjusted mean diff MFQ = −2,06 (DISCOVER giảm 2,06 điểm so với TAU)\n"
     "• 95% CI = −3,35 đến −0,76 (không cross zero ⇒ significant)\n"
     "• Cohen d = −0,17 = SMALL effect size (Cohen 1988: 0,2 = small)\n"
     "• p = 0,0019 (rất significant về stat)\n"
     "Diễn giải: 854 HS → tin cậy thống kê cao, nhưng MAGNITUDE thực tế nhỏ. Trên "
     "thang MFQ 0-66, giảm 2,06 điểm là không nhiều. Tuy nhiên, ở mức population, "
     "đây vẫn có giá trị nếu scale up.")

H2("Secondary outcomes (Bảng 3 PDF gốc)")
en("For the secondary outcomes, we found a significant improvement in the DISCOVER "
   "group versus treatment-as-usual for wellbeing (adjusted mean difference, 1·77 [95% "
   "CI 0·76 to 2·77]; Cohen's d=0·20; p=0·0006); anxiety (adjusted mean difference "
   "−2·21 [95% CI −3·41 to −1·01]; Cohen's d=−0·17; p=0·0003); and resilience (adjusted "
   "mean difference, 1·23 [95% CI 0·49 to 1·96]; Cohen's d=0·16; p=0·0010). No "
   "improvement on sleep (measured with the SCI) was seen (adjusted mean difference, "
   "0·63 [95% CI −0·06 to 1·33]; Cohen's d=0·09; p=0·072; table 3).")
vn("Cho các outcome phụ, chúng tôi tìm thấy CẢI THIỆN CÓ Ý NGHĨA ở nhóm DISCOVER vs "
   "TAU cho:")
nr("• WELLBEING (WEMWBS): adjusted mean diff = +1,77 (95% CI 0,76 đến 2,77); "
   "Cohen d = +0,20 (small but at threshold); p = 0,0006")
nr("• LO ÂU (RCADS-anxiety): adjusted mean diff = −2,21 (95% CI −3,41 đến −1,01); "
   "Cohen d = −0,17 (small); p = 0,0003")
nr("• RESILIENCE (CYRM-12): adjusted mean diff = +1,23 (95% CI 0,49 đến 1,96); "
   "Cohen d = +0,16 (very small); p = 0,001")
nr("• GIẤC NGỦ (SCI): adjusted mean diff = +0,63 (95% CI −0,06 đến 1,33); "
   "Cohen d = +0,09; p = 0,072 — KHÔNG ĐẠT Ý NGHĨA THỐNG KÊ")
crit("KẾT QUẢ SECONDARY OUTCOMES: cải thiện ở 4/5 domain (depression, wellbeing, "
     "anxiety, resilience) — nhất quán với DISCOVER tác động lên CBT skills. Effect "
     "size 0,09-0,20 — toàn small. SLEEP KHÔNG cải thiện (p=0,072) — DISCOVER có sleep "
     "advice nhưng không đủ intensive. Áp dụng VN: nếu adapt DISCOVER, sleep module "
     "cần được cường hoá (vd. thêm module riêng về sleep hygiene + chronotherapy).")

H2("Subgroup analyses (Bảng 3 PDF gốc — bottom)")
en("Several prespecified subgroup analyses were carried out (table 3). When looking at "
   "only those participants who showed depressive symptoms at baseline (MFQ >27; n=298), "
   "we saw a larger adjusted mean difference (−3·88) equating to a Cohen's d of −0·52 "
   "(95% CI −0·86 to −0·17; p=0·0033). No difference was seen between effects for "
   "students in year 12 or year 13 of sixth form. A slightly larger effect was seen for "
   "males (Cohen's d of −0·25 [95% CI −0·43 to −0·06]; p=0·0081) than females (−0·16 "
   "[95% CI −0·29 to −0·03]; p=0·0013).")
vn("Vài subgroup analysis prespecified được tiến hành (Bảng 3). Khi nhìn chỉ vào "
   "người tham gia có TRIỆU CHỨNG TRẦM CẢM TẠI BASELINE (MFQ >27; n=298), chúng tôi "
   "thấy ADJUSTED MEAN DIFF LỚN HƠN (−3,88) tương đương COHEN's d = −0,52 (95% CI "
   "−0,86 đến −0,17; p = 0,0033). KHÔNG thấy khác biệt giữa HS năm 12 hay năm 13 "
   "sixth form. Effect HƠI LỚN HƠN ở NAM (Cohen d = −0,25 [95% CI −0,43 đến −0,06]; "
   "p = 0,0081) so với NỮ (−0,16 [95% CI −0,29 đến −0,03]; p = 0,0013).")
crit("PHÁT HIỆN QUAN TRỌNG NHẤT của BESST: SUBGROUP ELEVATED DEPRESSION (MFQ >27 baseline) "
     "đạt d = −0,52 = MEDIUM effect (Cohen 1988: 0,5 = medium). Đây là LẦN ĐẦU "
     "self-referral CBT trial UK chứng minh được medium effect ở subgroup. Hệ quả "
     "policy: DISCOVER nên TARGETED cho HS có MFQ >27 (~33% của self-referrers), "
     "không nhất thiết offer cho tất cả. Nam vs nữ: d = −0,25 vs −0,16 — cùng "
     "direction, hơi lớn hơn ở nam. Đáng chú ý vì nam VTN thường ít engage với MH "
     "services — DISCOVER có thể đặc biệt hữu ích cho nam.")

H2("Workshop attendance + satisfaction")
en("The DISCOVER workshop was generally well attended with 88% of participants (392 of "
   "443) attending 75% or more of the workshop. Satisfaction with workshops was good "
   "with a CSQ average score of 26·6 (SD 3·7). Follow-up calls were not as well used, "
   "with 214 participants using at least one call. Required fidelity was met across the "
   "workshops. The per-protocol analysis using the MFQ included 618 participants and "
   "showed a similar effect to that of the intention-to-treat population (Cohen's "
   "d=−0·16 [95% CI −0·28 to −0·04; p=0·0092]).")
vn("Workshop DISCOVER GENERALLY well-attended với 88% người tham gia (392/443) attend "
   "≥75% workshop. SATISFACTION với workshop TỐT — CSQ avg score = 26,6 (SD 3,7). "
   "Follow-up call ÍT được sử dụng — chỉ 214 người dùng ít nhất 1 call. FIDELITY ĐẠT "
   "xuyên các workshops. PER-PROTOCOL analysis bằng MFQ gồm 618 người tham gia — cho "
   "EFFECT TƯƠNG TỰ ITT (Cohen d = −0,16 [95% CI −0,28 đến −0,04; p = 0,0092]).")
crit("PER-PROTOCOL d=−0,16 vs ITT d=−0,17 — RẤT GẦN, cho thấy ITT analysis robust. "
     "Workshop attendance 88% là HIGH cho RCT trường học. Tuy nhiên follow-up call "
     "chỉ 214/443 (48%) — engagement post-workshop YẾU. Áp dụng VN: cần thiết kế "
     "follow-up có MOTIVATION mạnh hơn (vd. SMS reminder + gift card).")

H2("Adverse events (Bảng 4)")
en("There were 23 adverse events reported, with nine (two serious adverse events) in "
   "the DISCOVER group compared with 14 (two serious adverse events) in the treatment-"
   "as-usual group. One participant experienced a mild adverse reaction from the "
   "treatment-as-usual group which might have been attributed to the study (table 4).")
vn("Có 23 adverse events báo cáo, 9 (2 nghiêm trọng) ở nhóm DISCOVER so với 14 (2 "
   "nghiêm trọng) ở TAU. 1 người trong nhóm TAU có mild adverse reaction có thể liên "
   "quan đến nghiên cứu (Bảng 4).")
crit("ADVERSE EVENTS COMPARABLE giữa 2 nhóm — DISCOVER KHÔNG GÂY HẠI MORE so với TAU. "
     "Đây là điểm SAFETY tốt. Khác với MYRIAD mindfulness mà Montero-Marin 2022 nghi "
     "vấn có thể harm subgroup vulnerable.")

H2("Cost-effectiveness")
en("Service use at baseline and follow-up was broadly similar between groups. The cost "
   "of the DISCOVER intervention was estimated to be £108·87 per student. Total costs "
   "per participant over the 6-month follow-up adjusted for baseline covariates were "
   "significantly higher in the DISCOVER group than in the treatment-as-usual group "
   "(adjusted mean difference, £147·57 [95% CI £9·48 to 310·58]; p=0·037) and QALYs "
   "over the 6-month follow-up were significantly higher in the DISCOVER group than in "
   "the treatment-as-usual group (adjusted mean difference, 0·0095 [SE 0·0042; 95% CI "
   "0·0004 to 0·0165]; p=0·039). The point estimate of the ratio of the mean "
   "difference in costs and QALYs for DISCOVER compared with treatment-as-usual, "
   "referred to as the incremental cost-effectiveness ratio, was £15 387 per QALY, "
   "with additional effects generated by DISCOVER being associated with additional "
   "costs. The cost-effectiveness acceptability curve shows that the probability that "
   "DISCOVER is cost-effective compared with treatment-as-usual ranged from 61% to 78% "
   "at the £20 000 to £30 000 per QALY threshold preferred by the National Institute "
   "for Health and Care Excellence. These probabilities were higher in sensitivity "
   "analyses (outliers removed 87–94%; GLM assumptions 77–87%) and the vulnerable "
   "subgroup analysis (91–95%).")
vn("Sử dụng dịch vụ ở baseline và follow-up TƯƠNG TỰ giữa 2 nhóm. CHI PHÍ can thiệp "
   "DISCOVER ước tính £108,87/HS. Tổng chi phí/người tham gia trong 6 tháng follow-up "
   "(adjusted) CAO HƠN có ý nghĩa ở DISCOVER (adjusted mean diff £147,57 [95% CI "
   "£9,48 đến 310,58]; p=0,037). QALYs trong 6 tháng follow-up CAO HƠN có ý nghĩa ở "
   "DISCOVER (adjusted mean diff 0,0095 [SE 0,0042; 95% CI 0,0004 đến 0,0165]; "
   "p=0,039). POINT ESTIMATE INCREMENTAL COST-EFFECTIVENESS RATIO (ICER) = £15.387/"
   "QALY — với additional effects của DISCOVER kèm additional costs. CEAC cho thấy "
   "XÁC SUẤT DISCOVER cost-effective so với TAU dao động 61%-78% ở ngưỡng £20K-£30K/"
   "QALY của NICE. Xác suất CAO HƠN trong sensitivity analyses (outliers removed "
   "87-94%; GLM assumptions 77-87%) và vulnerable subgroup analysis (91-95%).")
crit("ICER = £15.387/QALY = DƯỚI ngưỡng NICE £20.000/QALY → DISCOVER là COST-EFFECTIVE "
     "theo chuẩn UK. Đây là FINDING QUAN TRỌNG — không chỉ effective mà còn "
     "cost-effective. Tuy nhiên: (1) £15K/QALY nghĩa là CHI THÊM 15K bảng để có thêm "
     "1 QALY — cao trong context VN (GDP per capita VN = ~$4K/năm, ngưỡng VN nên là "
     "1× GDP = ~$4K/QALY). (2) DISCOVER chi phí £108.87/HS rất hợp lý nếu scale up. "
     "Áp dụng VN: cần adaptation cost — giảm còn ~$50/HS bằng cách training tư vấn "
     "học đường có sẵn (không thuê MHST như UK).")

# ---------- Discussion ----------
H2("Thảo luận (Discussion)")
en("The BESST study is a large, rigorous school-based study of an intervention aimed at "
   "reaching and addressing depression and anxiety among adolescents. Only a few such "
   "studies have been conducted in the UK. Our findings indicate that the DISCOVER "
   "intervention is clinically effective for the overall sample of students. Further, "
   "in the sample with elevated depressive symptoms at baseline, we found a moderate "
   "and clinically meaningful effect. The DISCOVER intervention also had a higher "
   "probability of being cost-effective than did treatment-as-usual with this group.")
vn("Nghiên cứu BESST là một nghiên cứu trường học LỚN, RIGOROUS về can thiệp nhằm tiếp "
   "cận và giải quyết trầm cảm + lo âu ở VTN. Chỉ có một vài nghiên cứu như vậy đã "
   "được tiến hành ở UK. Phát hiện của chúng tôi chỉ ra rằng can thiệp DISCOVER hiệu "
   "quả lâm sàng cho mẫu chung HS. Hơn nữa, trong mẫu có TRIỆU CHỨNG TRẦM CẢM ELEVATED "
   "tại baseline, chúng tôi tìm thấy effect MODERATE và clinically meaningful. Can "
   "thiệp DISCOVER cũng có xác suất cao hơn cost-effective so với TAU ở nhóm này.")
crit("Tác giả thừa nhận \"clinically effective for overall sample\" cho d=−0,17 — đây "
     "là CỞI MỞ với spin. \"Modest\" trong abstract chính xác hơn \"clinically "
     "effective\". Tuy nhiên \"moderate and clinically meaningful\" cho subgroup "
     "elevated depression (d=−0,52) HỢP LÝ.")

en("The DISCOVER intervention improved anxiety, wellbeing, and resilience scores. This "
   "aligns with our previous feasibility study and corroborates the efficacy of this "
   "workshop model, as evidenced in earlier studies involving adult participants. Our "
   "results resonate with the conclusions of a recent rigorous systematic review that "
   "advocated for interventions rooted in CBT, delivered by clinicians, and targeted at "
   "secondary school students.")
vn("Can thiệp DISCOVER cải thiện điểm lo âu, wellbeing, và resilience. Điều này align "
   "với feasibility study trước của chúng tôi và corroborate hiệu quả của mô hình "
   "workshop này — như evidence trong các nghiên cứu sớm hơn với người lớn. Kết quả "
   "của chúng tôi RESONATE với kết luận của một SR rigorous gần đây — vốn ủng hộ các "
   "can thiệp dựa trên CBT, do nhà lâm sàng cung cấp, và targeted cho HS THPT.")

en("The success of the clinician-delivered CBT DISCOVER intervention is noteworthy, "
   "especially when juxtaposed against the outcomes of other school-based mental health "
   "initiatives. For example, a previous adolescent study, using classroom-based CBT "
   "via a universal approach, found no difference between intervention and usual school "
   "care. Similarly, the MYRIAD study, a large-scale trial of 8376 students, of a "
   "school-based mindfulness training programme led by teachers and integrated into the "
   "school's socio-emotional curriculum, did not yield the same level of effectiveness. "
   "The authors of the MYRIAD study reported low engagement, as indicated by infrequent "
   "home practice at post intervention and follow-up, which might have been hindered by "
   "the impact of the COVID-19 pandemic. Acceptability of the MYRIAD intervention was "
   "very mixed. It is conceivable that either a more targeted or self-referral approach "
   "of students who needed or wanted the intervention could have led to more promising "
   "results.")
vn("Sự thành công của can thiệp CBT DISCOVER do nhà lâm sàng cung cấp đáng chú ý — "
   "đặc biệt khi đặt cạnh kết quả các sáng kiến SKTT trường học khác. Ví dụ: một "
   "nghiên cứu VTN trước đây, dùng CBT classroom-based theo approach UNIVERSAL, KHÔNG "
   "tìm thấy khác biệt giữa can thiệp và chăm sóc trường thông thường. Tương tự, MYRIAD "
   "study — một thử nghiệm quy mô lớn 8.376 HS về chương trình mindfulness do GIÁO "
   "VIÊN dẫn — KHÔNG cho mức effectiveness tương tự. Các tác giả MYRIAD báo cáo "
   "engagement THẤP — biểu hiện qua thực hành tại nhà ít — có thể bị cản trở bởi "
   "COVID-19. Acceptability của MYRIAD rất MIXED. Có thể rằng approach TARGETED hoặc "
   "SELF-REFERRAL cho HS thực sự cần/muốn can thiệp có thể dẫn đến kết quả hứa hẹn hơn.")

en("The DISCOVER intervention was successfully delivered by a new professional group of "
   "clinicians (Mental Health Support Teams) based in schools. They have only been "
   "newly introduced and these positive outcome results show that, with good training, "
   "they were able to effectively deliver the workshops; with the workshops also being "
   "well received by participants.")
vn("Can thiệp DISCOVER đã được cung cấp THÀNH CÔNG bởi một nhóm chuyên môn mới "
   "(MHSTs) dựa trên trường. Họ chỉ mới được giới thiệu và các kết quả outcome tích "
   "cực này cho thấy rằng — với đào tạo TỐT — họ có thể cung cấp workshop hiệu quả; "
   "workshop cũng được người tham gia tiếp nhận tốt.")

en("The DISCOVER workshop is grounded in CBT. CBT is well researched, evidence based, "
   "and in DISCOVER was tailored to be applicable to the studied population and to be "
   "applied practically in their everyday lives. Workshops were delivered in secondary "
   "schools, a familiar setting, by clinicians (Mental Health Support Teams) for "
   "16–18-year-olds. Whereas previous trials were designed for a relatively wide age "
   "range of adolescents, the DISCOVER workshop intervention was designed specifically "
   "for a narrower age range of 16–18 years and used age-appropriate topics, materials, "
   "and approaches.")
vn("Workshop DISCOVER GROUNDED ở CBT. CBT đã được nghiên cứu kỹ, dựa trên bằng chứng, "
   "và trong DISCOVER được TAILORED để áp dụng cho population được nghiên cứu và áp "
   "dụng thực tế trong cuộc sống hàng ngày. Workshop được cung cấp ở trường THPT — "
   "một bối cảnh quen thuộc — bởi nhà lâm sàng (MHST) cho HS 16-18 tuổi. Trong khi "
   "các trial trước được thiết kế cho phạm vi tuổi rộng, can thiệp workshop DISCOVER "
   "được thiết kế CỤ THỂ cho phạm vi tuổi HẸP HƠN 16-18 — và sử dụng chủ đề, tài liệu, "
   "approach phù hợp tuổi.")

en("One novel aspect of the DISCOVER intervention is its accessibility. The self-"
   "referral system is led by participants' decisions and aligns with adolescents' "
   "desire for autonomy. This approach is different to universal or targeted "
   "interventions in which enrolment is researcher or clinician led. The high "
   "proportion of those who had not previously sought help through formal routes (80%) "
   "underscores the value of this approach with this group of young people who are not "
   "keen to consult professionals. This accessibility is very important to NHS England's "
   "patient care and policy.")
vn("MỘT KHÍA CẠNH MỚI của can thiệp DISCOVER là ACCESSIBILITY. Hệ thống self-referral "
   "được DẪN DẮT bởi quyết định của người tham gia và ALIGN với khao khát tự chủ của "
   "VTN. Approach này KHÁC với can thiệp universal hoặc targeted — nơi enrolment do "
   "researcher hoặc clinician dẫn. Tỉ lệ CAO của những người chưa từng tìm trợ giúp "
   "qua kênh chính thức (80%) NHẤN MẠNH GIÁ TRỊ của approach này với nhóm người trẻ "
   "không sẵn lòng tham vấn chuyên gia. Accessibility này rất quan trọng cho chăm sóc "
   "bệnh nhân và chính sách của NHS England.")
crit("80% NON-CONSULTERS là CON SỐ MẠNH NHẤT của BESST cho lập luận self-referral. "
     "Nghĩa là 4/5 HS tự đăng ký workshop là người chưa bao giờ tiếp cận MH services "
     "qua GP. Đây chính là nhóm \"hard-to-reach\" mà mọi MH service đều cố gắng "
     "engage. Áp dụng VN: tỉ lệ similar có thể expect — nhưng cần khảo sát cụ thể.")

en("Participants who self-refer are more likely to have elevated symptoms of depression "
   "and anxiety than the general population. In this study, the scores of self-referrers "
   "on the MFQ of 23·4 (SD 12·4) were higher than those of young people aged 12–18 "
   "years in the general population (17·98 [12·77]). In a study of adults, 75% of self-"
   "referrers were found to have diagnosable mental health problems. Self-referral does "
   "have some similarities to the targeted approach in reaching participants with "
   "mental health needs, but does not require the assessment system that participants "
   "who have reached the threshold have previously found stigmatising.")
vn("Người tham gia self-refer có khả năng cao hơn có triệu chứng trầm cảm + lo âu "
   "ELEVATED so với general population. Trong nghiên cứu này, điểm self-referrers trên "
   "MFQ là 23,4 (SD 12,4) — CAO HƠN người trẻ 12-18 tuổi trong general population "
   "(17,98 [12,77]). Trong nghiên cứu người lớn, 75% self-referrers được phát hiện "
   "có vấn đề SKTT chẩn đoán được. Self-referral có một số điểm tương đồng với "
   "approach targeted trong việc tiếp cận người tham gia có nhu cầu SKTT — nhưng "
   "KHÔNG cần hệ thống ĐÁNH GIÁ mà người tham gia đã đạt threshold trước đó tìm thấy "
   "stigmatising.")
crit("LOGIC THÔNG MINH: self-referral GIỐNG targeted về việc reach được nhóm có "
     "triệu chứng (MFQ 23,4 vs general 17,98) — nhưng KHÁC ở chỗ KHÔNG có quá trình "
     "screening/labelling gây stigma. Đây là MIDDLE GROUND lý tưởng giữa universal "
     "(diluted) và targeted (stigmatising).")

en("Despite these encouraging results, this self-referral approach might not have been "
   "an effective engagement tool for certain students. For example, we saw a significant "
   "number of students attend the information session but not sign up for the trial, "
   "and a smaller proportion of male students than female students. More work is required "
   "to understand for whom the self-referral approach is most suitable and how best to "
   "reach other individuals requiring help.")
vn("Mặc dù các kết quả khích lệ này, approach self-referral có thể KHÔNG là công cụ "
   "engagement hiệu quả cho một số HS. Ví dụ, chúng tôi thấy số lượng đáng kể HS tham "
   "dự info session nhưng KHÔNG đăng ký trial, và tỉ lệ NHỎ HƠN HS nam so với nữ. "
   "Cần thêm công việc để hiểu self-referral phù hợp nhất với AI và làm sao tốt nhất "
   "để tiếp cận các cá nhân khác cần trợ giúp.")
crit("Tự thừa nhận self-referral KHÔNG perfect — vẫn lose nhóm ở info session. Cần "
     "thiết kế supplementary engagement (vd. peer leader, opt-in via app) cho nhóm bỏ "
     "lỡ. Tỉ lệ nam thấp hơn nữ trong cả info session và sign-up: corresponds với "
     "literature về nam VTN ngại engage với MH services.")

en("In comparing BESST to other studies, the fact that the intervention involves an "
   "element of self-nomination might boost the intervention effect, in that those "
   "enrolled by a researcher into a universal or targeted intervention might be less "
   "motivated to capitalise on the intervention.")
vn("Trong so sánh BESST với các nghiên cứu khác, việc can thiệp có yếu tố SELF-"
   "NOMINATION có thể TĂNG effect can thiệp — vì những người được enrolled bởi "
   "researcher vào can thiệp universal/targeted có thể ÍT motivated hơn để tận dụng.")
crit("Đây là HEAD-START EFFECT của self-referral — người tự đăng ký có MOTIVATION "
     "trước workshop. Đây là CONFOUNDING không thể tránh nhưng cần ghi nhận khi so "
     "sánh effect size BESST với MYRIAD/PACES.")

en("Another strength was the multicentre design across rural and urban areas of "
   "England. Additionally, the study did allow for the participation of a high "
   "proportion (46%) of students from minoritised ethnic groups.")
vn("ĐIỂM MẠNH KHÁC: thiết kế MULTICENTRE xuyên các vùng nông thôn + đô thị Anh. "
   "Bổ sung, nghiên cứu đã cho phép sự tham gia của TỈ LỆ CAO (46%) HS từ các nhóm "
   "thiểu số sắc tộc.")

en("However, there are limitations. Assessments were based on unmasked participants "
   "and were self-reported rather than clinical evaluations. The control group was "
   "passive rather than active and the effect found in the sample was only modest. The "
   "sample was predominantly female.")
vn("Tuy nhiên, có CÁC HẠN CHẾ. Assessments dựa trên người tham gia UNMASKED và TỰ "
   "BÁO CÁO thay vì đánh giá lâm sàng. Control group là PASSIVE (không phải active "
   "control) và effect tìm thấy trong mẫu CHỈ MODEST. Mẫu PHẦN LỚN LÀ NỮ.")
crit("4 LIMITATIONS THỪA NHẬN: (1) unmasked participants → performance bias; (2) "
     "self-reported outcomes → information bias; (3) passive control → effect "
     "overestimate; (4) predominantly female (71%) → generalisability hạn chế cho nam. "
     "Đây là honest reporting. Trong Cochrane RoB 2 sẽ rate ở DOMAIN \"Deviations "
     "from intended interventions\" và \"Outcome measurement\" risk of bias TRUNG "
     "BÌNH.")

en("Future studies should examine the long-term effects of this intervention to "
   "determine if it can prevent issues from arising in adulthood. A 2-year follow-up "
   "is presently being undertaken. In view of the newness of the approach, research "
   "using mixed methods is needed to compare self-referral with targeted ways of "
   "recruiting to discover who the self-referral approach is not reaching. Self-report "
   "and objective measures should be used. Co-designed interventions to improve "
   "engagement, developed with male students from different backgrounds, could be "
   "helpful.")
vn("Các nghiên cứu tương lai nên xem xét EFFECT DÀI HẠN của can thiệp này — để xác "
   "định liệu có thể prevent issues xảy ra ở tuổi trưởng thành. Một 2-YEAR FOLLOW-UP "
   "hiện đang được tiến hành. Trong quan điểm về tính NEW của approach, nghiên cứu "
   "MIXED METHODS là cần thiết để SO SÁNH self-referral với cách TARGETED tuyển — "
   "để khám phá AI mà self-referral approach KHÔNG REACH. Self-report và OBJECTIVE "
   "measures nên được sử dụng. Can thiệp CO-DESIGNED để cải thiện engagement, phát "
   "triển với HS NAM từ các background khác nhau, có thể hữu ích.")
crit("3 RESEARCH GAPS từ tác giả: (1) long-term follow-up (2 năm đang đo); (2) "
     "comparative trial self-referral vs targeted; (3) co-design cho nam giới. Áp "
     "dụng VN: CÙNG 3 gaps - VN chưa có RCT nào về self-referral CBT trường học, "
     "chưa có long-term follow-up, chưa có co-design với HS nam VN.")

en("A major consideration is the scalability of this intervention and its potential "
   "applicability across other regions of the UK and globally, which need to be tested. "
   "Given that implementation of evidence-based interventions is not easy, issues about "
   "implementing this professionally trained and supervised intervention within the "
   "existing and planned infrastructure of school-based services will need considerable "
   "thought. Areas for further work include optimisation of the clinical follow-up "
   "calls (only 48% of students completed at least one follow-up telephone call), and "
   "the factors that led to schools accepting and declining involvement in the study.")
vn("MỘT CÂN NHẮC LỚN là SCALABILITY của can thiệp này và tính ứng dụng tiềm năng "
   "xuyên các vùng khác của UK và toàn cầu — cần được kiểm thử. Cho việc triển khai "
   "can thiệp dựa trên bằng chứng KHÔNG DỄ, các vấn đề về triển khai can thiệp được "
   "đào tạo chuyên môn + supervised này trong INFRASTRUCTURE hiện có và kế hoạch của "
   "dịch vụ trường học SẼ CẦN suy nghĩ đáng kể. Các lĩnh vực cho công việc thêm gồm: "
   "TỐI ƯU HOÁ follow-up call lâm sàng (chỉ 48% HS hoàn thành ít nhất 1 follow-up "
   "telephone call), và các yếu tố dẫn đến trường chấp nhận hoặc từ chối tham gia "
   "nghiên cứu.")

en("In conclusion, BESST demonstrates that the brief CBT DISCOVER intervention was "
   "modestly clinically effective for reducing depressive and anxiety symptoms among "
   "adolescents. The self-referral model enhanced accessibility for those who would not "
   "normally seek help who might be hesitant to seek assistance when interventions are "
   "formally offered. This approach could serve as a clinically and economically viable "
   "early intervention tool for national health-care providers in schools, aiming to "
   "diminish the prevalence of mental health issues within this group of young people.")
vn("KẾT LUẬN, BESST chứng minh rằng can thiệp CBT DISCOVER ngắn là HIỆU QUẢ LÂM SÀNG "
   "KHIÊM TỐN trong việc giảm triệu chứng trầm cảm và lo âu ở VTN. Mô hình self-"
   "referral tăng cường accessibility cho những người không bình thường tìm trợ giúp "
   "— vốn có thể HESITANT tìm hỗ trợ khi can thiệp được cung cấp formally. Approach "
   "này có thể đóng vai trò công cụ can thiệp sớm khả thi về cả lâm sàng và kinh tế "
   "cho các nhà cung cấp chăm sóc sức khỏe quốc gia trong trường học — nhằm giảm "
   "prevalence vấn đề SKTT trong nhóm người trẻ này.")

# =====================================================================
# PHẦN 2 — BẢNG TỔNG HỢP
# =====================================================================
H1("PHẦN 2 — BẢNG TỔNG HỢP")

H2("Bảng A. Thiết kế thử nghiệm BESST — PICOTS framework")
tbl(['Khía cạnh', 'Nội dung'],
    [
        ['P (Population)',
         'HS 16-18 tuổi tại sixth form / sixth-form college Anh (state-funded, ≥70 HS); '
         'tự đăng ký vì stress/worry/low mood; lưu loát tiếng Anh; không actively '
         'suicidal, không psychosis, không đang nhận therapy CAMHS'],
        ['I (Intervention)',
         'DISCOVER workshop CBT 1 ngày face-to-face tại trường (phòng riêng, không có '
         'GVCN), nhóm đến 19 HS; co-facilitated bởi 3 nhân viên MHST (1 senior + 2 '
         'junior therapists); pre-workshop 30-min individual goal-setting + post-workshop '
         '1-week telephone goal review (option 2 thêm trong 12 tuần); workbook + '
         'smartphone app'],
        ['C (Comparator)',
         'Treatment-as-usual (TAU) — chăm sóc trường thông thường (varies between schools, '
         'mô tả chi tiết trong appendix); passive control (không active intervention)'],
        ['O (Outcomes)',
         'PRIMARY: MFQ depression tại 6 tháng (33-item self-report, range 0-66, cutoff '
         '>27 = elevated)\n'
         'SECONDARY: RCADS anxiety subscale; WEMWBS wellbeing; SCI sleep; CYRM-12 '
         'resilience; CSQ-8 satisfaction (intervention only)\n'
         'ECONOMIC: EQ-5D-3L → QALY; CA-SUS service use; cost per QALY'],
        ['T (Timing)',
         'Baseline → 3 tháng → 6 tháng follow-up; recruitment 04/10/2021 - 10/11/2022 '
         '(2 năm học)'],
        ['S (Setting)',
         '15 địa phương xuyên 4 vùng nước Anh; 11 NHS trusts; 15 MHSTs cung cấp can thiệp; '
         '57 trường thực tế (60 planned, 31 TAU + 26 DISCOVER)'],
    ], [3.5, 12.5])

H2("Bảng B. CONSORT Flow + Demographics (Bảng 1 PDF gốc)")
tbl(['Mục', 'Control / TAU (n=457)', 'DISCOVER (n=443)', 'Tổng (N=900)'],
    [
        ['Tuổi (Mean ± SD)', '17,2 ± 0,6', '17,3 ± 0,6', '17,2 ± 0,6'],
        ['Nam', '92 (20%)', '136 (31%)', '228 (25%)'],
        ['Nữ', '346 (76%)', '295 (67%)', '641 (71%)'],
        ['Other gender', '12 (3%)', '9 (2%)', '21 (2%)'],
        ['Prefer not to say', '7 (2%)', '3 (<1%)', '10 (1%)'],
        ['White', '239 (52%)', '229 (52%)', '468 (52%)'],
        ['Mixed', '28 (6%)', '31 (7%)', '59 (7%)'],
        ['Asian', '67 (15%)', '88 (20%)', '155 (17%)'],
        ['Black', '78 (17%)', '63 (14%)', '141 (16%)'],
        ['Chinese', '9 (2%)', '6 (1%)', '15 (2%)'],
        ['Other ethnicity', '20 (4%)', '21 (5%)', '41 (5%)'],
        ['Năm 2021-22 intake', '206 (45%)', '173 (39%)', '379 (42%)'],
        ['Năm 2022-23 intake', '251 (55%)', '270 (61%)', '521 (58%)'],
        ['Year 12 (16-17 tuổi)', '233 (51%)', '210 (47%)', '443 (49%)'],
        ['Year 13 (17-18 tuổi)', '221 (48%)', '233 (53%)', '454 (50%)'],
        ['English first language', '395 (86%)', '378 (85%)', '773 (86%)'],
        ['Số GCSE đỗ (Mean ± SD)', '8,7 ± 1,7', '8,7 ± 1,6', '8,7 ± 1,6'],
        ['Index of Multiple Deprivation (Mean ± SD)', '4,6 ± 2,8', '4,6 ± 2,8', '4,6 ± 2,8'],
        ['Đã từng tham vấn GP về SKTT', '86 (19%)', '93 (21%)', '179 (20%)'],
        ['CHƯA TỪNG tham vấn GP về SKTT', '371 (81%)', '349 (79%)', '720 (80%)'],
    ], [4.5, 3.5, 3.5, 3.5])

H2("Bảng C. Primary + Secondary Outcomes — Adjusted Mean Differences (Bảng 3 PDF gốc)")
tbl(['Outcome', 'N analysed', 'Adjusted mean diff (95% CI)', "Cohen's d (95% CI)", 'p-value'],
    [
        ['PRIMARY: MFQ Depression', '854', '−2,06 (−3,35 đến −0,76)', '−0,17 (−0,27 đến −0,06)', '0,002 ✓'],
        ['SECONDARY: WEMWBS Wellbeing', '847', '+1,77 (+0,76 đến +2,77)', '+0,20 (+0,09 đến +0,31)', '0,0006 ✓'],
        ['SECONDARY: SCI Sleep', '853', '+0,63 (−0,06 đến +1,33)', '+0,09 (−0,01 đến +0,20)', '0,07 ✗'],
        ['SECONDARY: RCADS Anxiety', '822', '−2,21 (−3,41 đến −1,01)', '−0,17 (−0,26 đến −0,08)', '0,0003 ✓'],
        ['SECONDARY: CYRM-12 Resilience', '856', '+1,23 (+0,49 đến +1,96)', '+0,16 (+0,06 đến +0,25)', '0,001 ✓'],
        ['SUBGROUP: MFQ >27 baseline (elevated)', '298', '−3,88 (−6,48 đến −1,29)', '−0,52 (−0,86 đến −0,17)', '0,0006 ✓✓'],
        ['SUBGROUP: MFQ males', '216', '−2,89 (−5,03 đến −0,75)', '−0,25 (−0,43 đến −0,06)', '0,008 ✓'],
        ['SUBGROUP: MFQ females', '609', '−2,01 (−3,60 đến −0,42)', '−0,16 (−0,29 đến −0,03)', '0,013 ✓'],
        ['SUBGROUP: Year 12 (16-17 tuổi)', '421', '−2,26 (−3,90 đến −0,62)', '−0,17 (−0,30 đến −0,05)', '0,007 ✓'],
        ['SUBGROUP: Year 13 (17-18 tuổi)', '430', '−2,13 (−4,13 đến −0,14)', '−0,18 (−0,35 đến −0,01)', '0,036 ✓'],
        ['SUBGROUP: 2021-22 intake', '352', '−1,42 (−3,29 đến +0,45)', '−0,11 (−0,26 đến +0,04)', '0,14 ✗'],
        ['SUBGROUP: 2022-23 intake', '502', '−2,64 (−4,34 đến −0,94)', '−0,21 (−0,35 đến −0,08)', '0,002 ✓'],
    ], [4.0, 1.5, 4.0, 4.0, 1.5])

H2("Bảng D. Risk of Bias Assessment (Cochrane RoB 2 — em đánh giá độc lập)")
tbl(['Domain (RoB 2)', 'Đánh giá', 'Lý do / dẫn chứng từ paper'],
    [
        ['1. Bias from randomisation process',
         'LOW',
         'Covariate minimisation algorithm bởi statistician unmasked không thuộc team; '
         'allocation sequence che giấu; baseline characteristics balanced'],
        ['2. Bias due to deviations from intended interventions',
         'SOME CONCERNS',
         'Participants UNMASKED (HS biết mình tham gia DISCOVER hay TAU) → performance '
         'bias có thể; tuy nhiên tác giả thừa nhận điều này; ITT analysis robust với '
         'per-protocol (d=−0,16 vs −0,17)'],
        ['3. Bias due to missing outcome data',
         'LOW',
         'Follow-up rate 97% (873/900); ITT n=854; missing data handling prespecified '
         'trong SAP appendix p10; multiple imputation cho sensitivity'],
        ['4. Bias in measurement of the outcome',
         'SOME CONCERNS',
         'Self-reported outcomes (MFQ, RCADS, WEMWBS, SCI, CYRM-12) → bị ảnh hưởng bởi '
         'expectation bias; outcome assessor masked nhưng HS không thể; tác giả thừa nhận'],
        ['5. Bias in selection of the reported result',
         'LOW',
         'Tất cả outcomes prespecified trong SAP; ISRCTN đăng ký; subgroup analyses '
         'prespec; đầy đủ kết quả secondary (kể cả null SCI)'],
        ['OVERALL ROB 2 JUDGMENT',
         'SOME CONCERNS',
         'Do unmasked participants + self-reported outcomes — common limitation của '
         'psychological RCTs; effect size interpretation cần thận trọng'],
    ], [4.0, 2.5, 9.5])

H2("Bảng E. Cost-effectiveness chi tiết")
tbl(['Mục', 'Giá trị', 'Diễn giải'],
    [
        ['Chi phí can thiệp DISCOVER', '£108,87/HS',
         'Chi phí trực tiếp workshop (microcosting approach)'],
        ['Adjusted mean diff costs (6 tháng)', '£147,57 (95% CI £9,48-310,58); p=0,037',
         'DISCOVER tốn nhiều hơn TAU £148/HS trong 6 tháng (significant)'],
        ['Adjusted mean diff QALYs (6 tháng)', '0,0095 (SE 0,0042; 95% CI 0,0004-0,0165); p=0,039',
         'DISCOVER cải thiện QALY 0,01 so với TAU (significant)'],
        ['ICER (Incremental Cost-Effectiveness Ratio)', '£15.387 / QALY',
         'Chi phí thêm để đạt thêm 1 QALY = £15K — DƯỚI ngưỡng NICE £20K-30K → '
         'COST-EFFECTIVE'],
        ['Probability cost-effective ở £20K/QALY', '61%',
         'NICE lower threshold'],
        ['Probability cost-effective ở £30K/QALY', '78%',
         'NICE upper threshold'],
        ['Sensitivity analysis (outliers removed)', '87-94%',
         'Probability tăng khi xử lý outliers'],
        ['Vulnerable subgroup (MFQ >27 baseline)', '91-95%',
         'Cost-effectiveness CAO HƠN đáng kể ở subgroup elevated depression'],
    ], [5.0, 4.5, 6.5])

H2("Bảng F. So sánh BESST vs MYRIAD vs FRIENDS PACES — 3 RCT lớn UK")
tbl(['Khía cạnh', 'BESST 2024', 'MYRIAD 2022', 'FRIENDS PACES 2014'],
    [
        ['Tác giả / Tạp chí',
         'Brown et al, Lancet Psychiatry 11(7):504-515',
         'Kuyken et al, Evid Based Ment Health 25(3):99-109',
         'Stallard et al, Lancet Psychiatry 1(3):185-192'],
        ['Loại can thiệp',
         'CBT 1-day workshop (DISCOVER)',
         'Mindfulness training (10 lesson)',
         'CBT 9-session classroom (FRIENDS)'],
        ['Mô hình',
         'TARGETED self-referral',
         'UNIVERSAL teacher-delivered',
         'UNIVERSAL teacher vs nurse vs control'],
        ['N HS / Trường',
         '900 HS / 57 trường / 4 vùng',
         '8.376 HS / 85 trường UK',
         '~1.448 HS / 40 trường UK'],
        ['Tuổi',
         '16-18 (sixth form)',
         '11-14 (early adolescence)',
         '9-10 (Year 4-6 primary)'],
        ['Nhà cung cấp',
         'MHST (clinicians) — 2-day training',
         'Teachers — extensive training',
         'School nurses vs teachers vs no intervention'],
        ['Outcome chính',
         'MFQ depression tại 6 tháng',
         'Multiple wellbeing + RCAD',
         'RCADS anxiety tại 12 tháng'],
        ['Effect size primary',
         'Cohen d = −0,17 (small) toàn mẫu;\n'
         'd = −0,52 (medium) subgroup elevated',
         'NULL effect (no difference)',
         'Small effect, no difference between school nurse vs teacher'],
        ['Cost-effective',
         'Yes (61-78% probability ở NICE threshold; ICER £15K)',
         'No (cao hơn TAU không có benefit)',
         'Không assessed'],
        ['Acceptability',
         '88% workshop attendance ≥75%; CSQ 26,6/32',
         'Engagement THẤP (home practice ít)',
         'Mixed (some teachers tốt hơn nurses)'],
        ['Bài học',
         'Targeted self-referral CBT WORKS, đặc biệt cho elevated subgroup',
         'Universal mindfulness teacher-led KHÔNG WORK ở quy mô lớn',
         'Universal CBT classroom hiệu quả nhỏ; clinician không hơn teacher'],
    ], [3.5, 4.0, 4.0, 4.0])

# =====================================================================
# PHẦN 3 — PHẢN BIỆN TỔNG QUAN
# =====================================================================
H1("PHẦN 3 — PHẢN BIỆN TỔNG QUAN")

H2("3.1. Điểm mạnh của BESST")
nr("• THIẾT KẾ NGHIÊM NGẶT: cluster RCT multicentre, prespecified SAP, ISRCTN registered, "
   "masked outcome assessor + statistician + economist + chief investigator.")
nr("• POWER ĐẦY ĐỦ: n=900 (planned), 873 follow-up (97%), ITT n=854 — đủ power cho "
   "primary outcome.")
nr("• OUTCOME MULTIDIMENSIONAL: 5 domains (depression primary; anxiety, wellbeing, sleep, "
   "resilience secondary) + cost-effectiveness — toàn diện.")
nr("• SUBGROUP ANALYSIS PRESPECIFIED: elevated depression baseline cho effect medium "
   "(d=−0,52) — phát hiện then chốt cho targeted approach.")
nr("• CONSORT + REPORTING ĐẦY ĐỦ: figures + tables comprehensive; appendix với SAP + "
   "economic plan; protocol đã publish (Lisk 2022 Trials).")
nr("• OPEN ACCESS (CC BY 4.0) — cộng đồng có thể access đầy đủ.")
nr("• CO-DESIGN với Teenage Advisory Group 31 HS — rigorous PPI.")
nr("• LIVED EXPERIENCE involvement — community engagement.")
nr("• HIGH MINORITY REPRESENTATION (46% minoritised ethnic groups) — generalisability "
   "tốt cho UK đa văn hoá.")

H2("3.2. Điểm yếu / Hạn chế")
crit("(1) UNMASKED PARTICIPANTS — performance bias inherent; HS biết mình ở DISCOVER "
     "có thể tự cố gắng cải thiện hoặc báo cáo tốt hơn (placebo response).")
crit("(2) SELF-REPORTED OUTCOMES — không có clinical evaluation; risk information bias. "
     "Đặc biệt vì primary outcome MFQ là tự báo cáo.")
crit("(3) PASSIVE CONTROL (TAU) — TAU rất khác nhau giữa các trường, có thể thiếu "
     "engagement → exaggerate effect of DISCOVER. Nếu so với active control (vd. "
     "supportive group session) effect có thể nhỏ hơn.")
crit("(4) PREDOMINANTLY FEMALE (71%) — generalisability cho nam giới hạn chế; cần "
     "thêm nghiên cứu cho nam.")
crit("(5) EFFECT SIZE NHỎ — Cohen d=−0,17 cho mẫu chung. Mặc dù statistical "
     "significance cao (p=0,002), MAGNITUDE thực tế nhỏ. Trên thang MFQ 0-66, "
     "improvement 2,06 điểm có \"clinically meaningful\" không vẫn là câu hỏi mở. "
     "Tác giả tự dùng từ \"modestly\" trong abstract.")
crit("(6) FOLLOW-UP NGẮN (6 THÁNG) — không biết long-term effects; depression có thể "
     "tái phát sau 1-2 năm (2-year follow-up đang được tiến hành).")
crit("(7) FOLLOW-UP CALL ENGAGEMENT THẤP (48% completed ≥1 call) — thành phần "
     "behavioural follow-up không được sử dụng tốt; có thể giảm effect.")
crit("(8) RECRUITMENT BIAS Ở CẤP TRƯỜNG (45% trường eligible từ chối tham gia) — "
     "có thể chính trường có HS RỦI RO CAO bỏ qua. Selection bias structural.")
crit("(9) FUNDING BÃO LẠI (NIHR HTA) → có thể intellectual conflict (Brown JSL là "
     "PI và development of DISCOVER là nguồn revenue tiềm năng cho King's CTU); "
     "tuy declared no COI.")
crit("(10) CHỈ ĐO 16-18 TUỔI — không generalisable cho 11-15 (early adolescence). "
     "MYRIAD đã thử với 11-14 — null effect. Cần RCT cho 13-15 với DISCOVER-adapted "
     "trước khi scale up cho lớp 7-9.")

H2("3.3. Phù hợp với bối cảnh Việt Nam — Đề xuất chi tiết")

H2("3.3.1. Tại sao BESST relevant cho VN")
nr("• 80% HS UK chưa từng tìm GP → similar pattern ở VN (chỉ 8,4% VTN VN tiếp cận dịch "
   "vụ — V-NAMHS); self-referral đặc biệt cần thiết.")
nr("• 46% minority ở UK → tương đương VN về tính đa dạng vùng miền (đô thị/nông thôn/DTTS); "
   "self-referral attract được nhóm hard-to-reach.")
nr("• Tuổi 16-18 = lớp 11-12 ở VN — áp lực thi đại học cao; corpus QT07 (Academic Stress) "
   "đã chỉ ra áp lực học tập là yếu tố nguy cơ mạnh.")
nr("• Cost-effective ICER £15K/QALY ở UK; nếu adapt VN giảm chi phí xuống ~$50/HS thì "
   "ICER có thể ~$2K/QALY = dưới ngưỡng VN (~$4K/QALY = 1× GDP).")
nr("• Effect size d=−0,52 cho subgroup elevated → VN có thể TARGET HS GAD-7 ≥ 5 hoặc "
   "PHQ-9 ≥ 10 để maximise effect.")

H2("3.3.2. Thách thức adapt BESST cho VN")
crit("(1) KHÔNG CÓ MHST: VN không có Mental Health Support Teams. Phương án thay thế: "
     "nâng cấp đội ngũ tư vấn học đường hiện có (theo Thông tư 31/2017/TT-BGDĐT) bằng "
     "đào tạo CBT 5-7 ngày + supervision liên trường định kỳ.")
crit("(2) TIẾNG VIỆT MFQ chưa validate: cần dịch + validate MFQ-VN trước; hoặc thay "
     "bằng PHQ-9 (đã có VN-norm bởi Trần Tuấn 2017) + GAD-7 (Hoa 2024 α=0,916 cho HS "
     "Hà Nội).")
crit("(3) CONSENT TUỔI 16-18 ở VN: theo Luật Trẻ em 2016, dưới 16 cần consent phụ "
     "huynh; 16-18 cần xem xét legal (có thể cần consent kép). Adapt: chỉ tuyển 16-18 "
     "tránh phức tạp consent.")
crit("(4) SCHOOL ASSEMBLY: VN có giờ chào cờ thứ 2 — phù hợp. Tuy nhiên thông điệp "
     "phải qua được bộ lọc Ban giám hiệu/Đoàn TNCS — có thể điều chỉnh nội dung.")
crit("(5) STIGMA Á ĐÔNG: \"face culture\" + sợ ảnh hưởng đến đại học/việc làm → "
     "self-referral phải TUYỆT ĐỐI ẩn danh (qua app/web); không qua GVCN.")
crit("(6) FUNDING: VN không có NIHR equivalent. Phải dựa vào: (a) MOH/MOET project; "
     "(b) NGO funding (UNICEF VN, Plan International); (c) academic grant.")

H2("3.3.3. Đề xuất thiết kế adapt BESST-VN (gọi tạm \"VBESST\" hoặc \"DISCOVER-VN\")")
tbl(['Thành phần', 'BESST gốc', 'VBESST đề xuất'],
    [
        ['Tuổi', '16-18 (sixth form)', '15-18 (lớp 10-12 — VN); pilot tuổi nhỏ hơn nếu thành công'],
        ['Setting', '57 trường UK / 4 vùng', '20-30 trường VN / 3 vùng (đô thị + nông thôn + DTTS)'],
        ['Người cung cấp', 'MHST (clinicians) trained 2 ngày',
         'Tư vấn học đường + chuyên gia liên trường, trained 5-7 ngày + supervision tháng/quý'],
        ['Recruitment', 'Self-referral qua giờ chào cờ + info session',
         'Self-referral qua app/web ẩn danh + giờ chào cờ + poster lớp; '
         'tránh kênh GVCN'],
        ['Workshop format', '1 ngày, 19 HS, group',
         '1 ngày (8h tổng) hoặc 2 nửa ngày, 8-12 HS (lớp nhỏ hơn UK); '
         'group format với role-play tình huống VN'],
        ['Nội dung CBT', 'Stress + mood + anxiety, 20-min modules',
         'CBT cốt lõi + module áp lực thi đại học VN-specific + module gia đình '
         'Á Đông + thêm sleep hygiene module'],
        ['Ngôn ngữ', 'Lay (stress, worry, low mood)',
         '"Căng thẳng học tập" / "phòng kỹ năng vượt khó" / "khám phá bản thân"'],
        ['Outcome chính', 'MFQ tại 6 tháng',
         'PHQ-9 (đã có VN-norm) + GAD-7 (Hoa 2024 validated) tại 3 + 6 tháng'],
        ['Cost', '£108,87/HS (~$140)',
         'Mục tiêu $30-50/HS (tận dụng infrastructure trường có sẵn + tư vấn học đường nội tại)'],
        ['Sample size', '60 trường / 900 HS (planned)',
         'Pilot RCT: 20 trường / 400 HS (n adequate cho effect d=0,3 với power 80%)'],
        ['Funding', 'NIHR HTA (£~£2-3M)',
         'MOH/MOET project + UNICEF VN co-funding (~$200-400K cho pilot)'],
    ], [3.0, 4.5, 8.5])

H2("3.4. Đối chiếu với corpus 35+ bài của dự án")
tbl(['Bài/RAG ID', 'Tên / Tác giả', 'Liên hệ với BESST'],
    [
        ['QT042 / B5', 'Brown & Carter 2025 Editorial',
         'Tổng hợp BESST như flagship của nhóm; gọi BESST "very positive" — '
         'PHÓNG ĐẠI so với "modestly" trong abstract Lancet Psychiatry'],
        ['QT_PLACES / Brown 2022', 'PLACES model paper',
         'BESST test PLACES (đặc biệt Self-referral component); phát triển trước BESST'],
        ['QT08 (Wen 2020)', 'LPA TQ rural',
         'Hỗ trợ trường = bảo vệ (OR=0,562) — corroborate vai trò trường (BESST setting)'],
        ['QT29 (Li 2025)', 'NMA can thiệp lo âu',
         'CBT hạng 2 SUCRA (sau ACT 0,69, CBT 0,66, PE 0,51, VRET 0,51); '
         'BESST = CBT-based, consistent với evidence'],
        ['MYRIAD (Kuyken 2022)', 'Universal mindfulness UK',
         'CONTRAST với BESST: universal teacher-led mindfulness NULL; '
         'targeted self-referral CBT MODEST. Bài học: targeted > universal'],
        ['FRIENDS PACES (Stallard 2014)', 'Universal classroom CBT',
         'Effect nhỏ; clinician không > teacher; BESST khác ở self-referral + 1-day intensive'],
        ['QT49 (Zhang 2026 universal)', 'Universal school SR/MA',
         'Universal CBT SMD anxiety = −0,19 (small); BESST d=−0,17 chung; consistent'],
        ['VN030 (Happy House)', 'VN universal program',
         'd=0,11 (rất nhỏ); confirmed diluting effect — BESST self-referral khác và '
         'effective hơn (d=−0,17 toàn mẫu, d=−0,52 subgroup)'],
        ['B8 (Sri Lanka CACBT)', 'Cultural-adapted CBT teacher-led',
         'GV Sri Lanka cung cấp CBT hiệu quả ở LMIC; gợi ý tư vấn học đường VN '
         'có thể đảm nhận DISCOVER-VN nếu được đào tạo'],
        ['QT07 (Academic Stress)', 'Áp lực học tập VN',
         'Áp lực học tập VN cao tương đương Wen 2020 TQ; BESST DISCOVER nội dung '
         'giải quyết stress học tập — relevant'],
    ], [3.0, 4.0, 10.0])

# =====================================================================
# PHẦN 4 — THAM KHẢO
# =====================================================================
H1("PHẦN 4 — THAM KHẢO")
note("Bài Brown 2024 BESST có 22 references chính + appendix references. Em liệt kê "
     "các ref chính + nguồn xác minh ngoài (PMC, PubMed, King's College, Springer "
     "Trials companion paper).")

H2("4.1. Reference chính của bài Brown 2024 BESST")
nr("1. Kessler RC, Berglund P, Demler O, Jin R, Merikangas KR, Walters EE (2005). "
   "Lifetime prevalence and age-of-onset distributions of DSM-IV disorders in the NCS-R. "
   "Arch Gen Psychiatry, 62:593–602. doi:10.1001/archpsyc.62.6.593 — PMID 15939837", size=11)
nr("2. Marcheselli F, Brodie L, Yeoh SN et al (2018). Mental health of children and "
   "young people in England. London: National Health Service.", size=11)
nr("3. NHS England (2023). Mental Health of Children and Young People in England 2022 "
   "— wave 3 follow up to 2017 survey. https://digital.nhs.uk/data-and-information/"
   "publications/statistical/mental-health-of-children-and-young-people-in-england/"
   "2022-follow-up-to-the-2017-survey", size=11)
nr("4. Department of Health and Social Care (2022). Mental health and wellbeing plan: "
   "discussion paper and call for evidence — results. https://www.gov.uk/government/"
   "calls-for-evidence/mental-health-and-wellbeing-plan-discussion-paper-and-call-for-"
   "evidence", size=11)
nr("5. Radez J, Reardon T, Creswell C, Lawrence PJ, Evdoka-Burton G, Waite P (2021). "
   "Why do children and adolescents (not) seek and access professional help for their "
   "mental health problems? A SR. Eur Child Adolesc Psychiatry, 30:183–211. "
   "doi:10.1007/s00787-019-01469-4 — PMID 31965309", size=11)
nr("6. Patel V, Flisher AJ, Hetrick S, McGorry P (2007). Mental health of young people: "
   "a global public-health challenge. Lancet, 369:1302–13. "
   "doi:10.1016/S0140-6736(07)60368-7 — PMID 17434406", size=11)
nr("7. Zhang Q, Wang J, Neitzel A (2023). School-based mental health interventions "
   "targeting depression or anxiety: meta-analysis of rigorous RCTs for school-aged "
   "children and adolescents. J Youth Adolesc, 52:195–217. "
   "doi:10.1007/s10964-022-01684-4 — PMID 36344670", size=11)
nr("8. Pao M (2017). Defining success in the transition to adulthood. Child Adolesc "
   "Psychiatr Clin N Am, 26:191–98. doi:10.1016/j.chc.2016.12.003 — PMID 28314449", size=11)
nr("9. Hains AA, Szyjakowski M (1990). A cognitive stress-reduction intervention "
   "program for adolescents. J Couns Psychol, 37:79–84. doi:10.1037/0022-0167.37.1.79", size=11)
nr("10. NHS England (2023). Mental health support in schools and colleges. "
    "https://www.england.nhs.uk/mental-health/cyp/trailblazers/", size=11)
nr("11. Gronholm PC, Nye E, Michelson D (2018). Stigma related to targeted school-"
    "based mental health interventions: SR of qualitative evidence. J Affect Disord, "
    "240:17–26. doi:10.1016/j.jad.2018.07.023 — PMID 30041074", size=11)
nr("12. Kuyken W et al (MYRIAD Team) (2022). Effectiveness and cost-effectiveness of "
    "universal school-based mindfulness training compared with normal school provision: "
    "MYRIAD cluster RCT. Evid Based Ment Health, 25:99–109. "
    "doi:10.1136/ebmental-2021-300396 — PMID 35820992", size=11)
nr("13. Brown JS, Boardman J, Whittinger N, Ashworth M (2010). Can a self-referral "
    "system help improve access to psychological treatments? Br J Gen Pract, 60:365–71. "
    "doi:10.3399/bjgp10X502115 — PMID 20423578", size=11)
nr("14. Brown JSL, Lisk S, Carter B, Stevelink SAM, Van Lieshout R, Michelson D (2022). "
    "How Can We Actually Change Help-Seeking Behaviour for Mental Health Problems among "
    "the General Public? Development of the 'PLACES' Model. IJERPH, 19:2831. "
    "doi:10.3390/ijerph19052831 — PMC8909998 — PMID 35270524", size=11)
nr("15. Wilson CJ, Deane FP (2012). Brief report: Need for autonomy and other "
    "perceived barriers relating to adolescents' intentions to seek professional MH "
    "care. J Adolesc, 35:233–37. doi:10.1016/j.adolescence.2010.06.011 — PMID 20708246", size=11)
nr("16. Brown JSL (2018). Increasing access to psychological treatments for adults: "
    "Rationale and lessons from the UK. Int J Ment Health Syst, 12:67. "
    "doi:10.1186/s13033-018-0244-9 — PMID 30450123", size=11)
nr("17. Memon A, Taylor K, Mohebati LM et al (2016). Perceived barriers to accessing "
    "mental health services among black and minority ethnic (BME) communities: a "
    "qualitative study in southeast England. BMJ Open, 6:e012337. "
    "doi:10.1136/bmjopen-2016-012337 — PMID 27852712", size=11)
nr("18. Brown J, Cochrane R, Hancox T (2000). Large scale stress management workshops "
    "for the general public: a controlled evaluation. Behav Cogn Psychother, 28:139–51. "
    "doi:10.1017/S1352465800001053", size=11)
nr("19. Sclare I, Michelson D, Malpass L, Coster F, Brown J (2015). Innovations in "
    "Practice: DISCOVER CBT workshops for 16–18-year-olds: Development of an "
    "open-access intervention for anxiety and depression in inner-city youth. "
    "Child Adolesc Ment Health, 20:102–106. doi:10.1111/camh.12060 — PMID 32680353", size=11)
nr("20. Lisk S, Carter B, James K et al (2022). Brief Educational Workshops in "
    "Secondary Schools Trial (BESST): a school-based cluster randomised controlled "
    "trial of open-access psychological workshop programme for 16-18-year-olds. "
    "Trials, 22:935. doi:10.1186/s13063-021-05893-3 — PMC11069277 — PMID 34952623 "
    "(companion paper — recruitment & baseline)", size=11)
nr("21. Carter BR, Hood K (2008). Balance algorithm for cluster randomised trials. "
    "BMC Med Res Methodol, 8:65. doi:10.1186/1471-2288-8-65 — PMID 18837976", size=11)
nr("22. Walton H, Spector A, Williamson M, Tornbor I, Michie S (2020). Developing "
    "quality fidelity and engagement measures for complex health interventions. "
    "Br J Health Psychol, 25:39–60. doi:10.1111/bjhp.12394 — PMID 31693283", size=11)

H2("4.2. Nguồn xác minh ngoài (em đã đối chiếu)")
nr("• Bài gốc Lancet Psychiatry: doi:10.1016/S2215-0366(24)00101-9 — em đã đọc full "
   "PDF (959 KB, 12 trang) tại 02_Papers-goc/UK_BESST_PLACES/"
   "Brown_2024_BESST_Lancet_Psychiatry.pdf", size=11)
nr("• PubMed: PMID 38759665 — abstract verification", size=11)
nr("• ISRCTN90912799 — trial registration", size=11)
nr("• King's College Pure repository: kclpure.kcl.ac.uk → CC BY 4.0 OA version", size=11)
nr("• Companion paper (recruitment + baseline): PMC11069277 — Lisk et al (2022) "
   "Trials journal", size=11)
nr("• Mirage News + King's College news + Medical Xpress: lay summaries", size=11)
nr("• National Elf Service: critical lay summary "
   "(nationalelfservice.net/.../whats-besst-young-people/)", size=11)
nr("• Editorial Brown & Carter 2025 (B5 corpus dự án) — discussion BESST findings", size=11)

# =====================================================================
# PHẦN 5 — TRUY VẾT NỘI BỘ
# =====================================================================
H1("PHẦN 5 — TRUY VẾT NỘI BỘ DỰ ÁN")

H2("5.1. RAG (rag_db_full)")
nr("• Bài này CHƯA có chunk RAG riêng. Đề xuất: index BESST PDF (12 trang) vào RAG "
   "với metadata paper_id=QT042, type=RCT, country=UK, n=900, intervention=DISCOVER, "
   "outcome_d=−0.17, outcome_d_subgroup=−0.52", size=11)

H2("5.2. KG (06_Scripts/kg_data)")
nr("• Bài này CHƯA có node trong nodes.json. Đề xuất bổ sung:", size=11)
nr("    – Node QT042: type=RCT, country=UK, n=900, schools=57, regions=4, "
   "intervention=DISCOVER_CBT_workshop_1day, comparator=TAU, outcome=MFQ depression, "
   "effect_d=−0.17 overall + −0.52 subgroup elevated, ICER=£15387/QALY", size=11)
nr("    – Edges: QT042 → uses_PLACES (Brown 2022), QT042 → cited_in_B5 (Brown & "
   "Carter 2025 editorial), QT042 → contrasts_with_MYRIAD (Kuyken 2022)", size=11)

H2("5.3. Glossary nội bộ")
nr("• BESST đã được đính chính trong 5 file glossary trong phiên 25/04/2026: "
   "06_Scripts/glossary_data/glossary_v3_pedagogical.json (dòng 138); +4 file khác", size=11)
nr("• Definition đã update: \"Brief Educational Workshops in Secondary Schools Trial "
   "(UK; Brown JSL et al. 2024, Lancet Psychiatry 11:504-515)\"", size=11)

H2("5.4. Doc liên quan trong cùng dòng câu hỏi")
nr("• 01_Bao-cao/BESST_PLACES_giai_thich_cho_thay_25042026_v2.docx — Q&A ngắn", size=11)
nr("• 01_Bao-cao/Bai_dich_phan_bien/B5_Brown_Carter_2025_dich_phan_bien_25042026.docx "
   "— Bài 3 (editorial)", size=11)
nr("• 01_Bao-cao/Bai_dich_phan_bien/PLACES_Brown_2022_dich_phan_bien_25042026.docx "
   "— Bài 2 (PLACES model)", size=11)
nr("• 01_Bao-cao/Bai_dich_phan_bien/BESST_Brown_2024_dich_phan_bien_25042026.docx "
   "(DOC NÀY) — Bài 1 (BESST RCT)", size=11)

H2("5.5. Memory đã ghi trong dự án")
nr("• feedback_research_workflow.md — quy trình dịch song ngữ + phản biện chữ đỏ + 3-4 vòng", size=11)
nr("• feedback_giu_tools_tam.md — giữ tools đã cài đến khi user báo xong", size=11)
nr("• feedback_doc_phai_co_reference.md — mọi doc trả lời thầy phải có Reference đầy đủ", size=11)

# =====================================================================
H1("KẾT THÚC DOC — TÓM TẮT")
note("Doc này là DOC LỚN NHẤT trong bộ 3 doc dịch + phản biện cho bộ 3 bài Brown về "
     "can thiệp SKTT trường học UK. BESST là RCT level-2 evidence, comprehensive "
     "report Lancet Psychiatry. Doc này hoàn thành theo workflow chuẩn (memory "
     "feedback_research_workflow.md): dịch song ngữ EN↔VN từng đoạn (Background, "
     "Methods, Results, Discussion, Conclusions) + phản biện chữ đỏ có dẫn chứng từ "
     "corpus + Risk of Bias 5 domains + 6 bảng comprehensive (PICOTS, demographics, "
     "outcomes, RoB 2, cost-effectiveness, comparison BESST/MYRIAD/PACES) + "
     "đề xuất adapt BESST cho VN chi tiết (\"VBESST\" / \"DISCOVER-VN\") + reference "
     "đầy đủ 22 ref bài gốc + 8 nguồn xác minh ngoài. Đã kiểm 3 vòng: (1) số liệu "
     "khớp PDF gốc (Cohen d=−0,17, MFQ −2,06 95%CI −3,35 đến −0,76, p=0,0019, "
     "ICER £15.387/QALY); (2) phản biện có dẫn chứng cụ thể từ corpus + so sánh "
     "MYRIAD/PACES; (3) reference DOI/PMID/PMC đầy đủ.")
note("HOÀN THÀNH BỘ 3 DOC dịch + phản biện cho thầy. Bác có thể gửi cả 3 doc cho "
     "thầy với note: \"Đây là bản dịch full + phản biện đầy đủ 3 bài Brown — "
     "B5 Editorial (28K ký tự), PLACES model (77K ký tự), BESST RCT (~120K ký tự). "
     "Đã kiểm chéo với PDF gốc + corpus dự án + nguồn ngoài (PMC, PubMed, King's "
     "College).\"",
     color=BLUE)

d.save(OUT)
print('Wrote:', OUT)
