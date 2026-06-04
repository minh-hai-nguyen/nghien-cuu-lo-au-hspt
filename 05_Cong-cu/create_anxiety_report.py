# -*- coding: utf-8 -*-
"""Báo cáo riêng: Rối loạn lo âu ở học sinh THCS — Việt Nam và Thế giới"""
import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

BASE = r"c:\Users\OS\OneDrive\read_books\Lo-au"
doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)

doc.add_heading('RỐI LOẠN LO ÂU Ở HỌC SINH THCS', level=0)
doc.add_heading('Tổng quan tình hình tại Việt Nam và Thế giới', level=1)
p = doc.add_paragraph()
run = p.add_run('Báo cáo này tổng hợp bằng chứng khoa học về rối loạn lo âu ở học sinh trung học cơ sở (10-15 tuổi) từ các nguồn uy tín tại Việt Nam, Đông Nam Á và thế giới.')
run.font.size = Pt(11)
doc.add_paragraph()

# PART 1: GLOBAL OVERVIEW
doc.add_heading('PHẦN 1: TỔNG QUAN TOÀN CẦU', level=1)

doc.add_heading('1.1. Tỷ lệ lo âu ở trẻ em và thanh thiếu niên trên thế giới', level=2)

global_data = [
    ['Nguồn', 'Phạm vi', 'Nhóm tuổi', 'Tỷ lệ lo âu', 'Ghi chú'],
    ['Polanczyk et al.\n(2015)\nJ Child Psychol\nPsychiatry, Q1', '41 nghiên cứu\n27 quốc gia', 'Trẻ em\n& TN (thanh thiếu niên)', '6.5%\n(chẩn đoán)', 'Meta-analysis CHUẨN VÀNG về tỷ lệ rối loạn tâm thần ở trẻ em [2].\nChỉ bao gồm nghiên cứu sử dụng tiêu chí CHẨN ĐOÁN lâm sàng.\nTỷ lệ 6.5% thấp hơn nhiều so với sàng lọc (20-70% ở 11 bài gốc).\nChênh lệch cho thấy phần lớn triệu chứng sàng lọc không đạt ngưỡng chẩn đoán.\nĐược trích dẫn >3,000 lần — nguồn tham khảo hàng đầu.'],
    ['Racine et al.\n(2021)\nJAMA Pediatrics\nQ1, IF~26', '29 nghiên cứu\n>80,000 TN', 'Trẻ em\n& TN\nCOVID-19', '20.5%\n(sàng lọc)', 'Meta-analysis QUAN TRỌNG NHẤT về SKTT (sức khỏe tâm thần) TN trong COVID-19 [1].\nTỷ lệ lo âu gấp ĐÔI so với trước đại dịch.\nNữ và thanh thiếu niên lớn tuổi bị ảnh hưởng nặng hơn.\nThu thập càng muộn trong đại dịch, tỷ lệ càng cao.\nXu (2021) báo cáo 9.89% — thấp hơn nhiều, có thể do thu thập rất sớm (2/2020).\nĐược trích dẫn >2,000 lần — tạp chí IF~26.'],
    ['GBD Study\n(2021)\nTransl Psychiatry\nQ1', '204 quốc gia\n1990-2021', '10-14\ntuổi', '4.1%\n(GBD)', 'Phân tích TOÀN DIỆN NHẤT: 204 quốc gia, 31 năm [3].\nNhóm 10-14 tuổi (THCS) có gánh nặng DALY (Disability-Adjusted Life Years — Số năm sống điều chỉnh theo khuyết tật) tâm thần CAO NHẤT: 16.3%.\nĐiều này có nghĩa ở lứa tuổi THCS, rối loạn tâm thần gây mất sức khỏe nhiều nhất.\nLo âu tăng >50% từ 1990 đến 2021 ở nhóm này.\nTăng mạnh nhất giai đoạn 2019-2021 do COVID-19 [3].\nViệt Nam nằm trong phân tích — tỷ lệ ở mức trung bình khu vực.'],
    ['GBD Study\n(2021)', '204 quốc gia', '15-19\ntuổi', '5.5%\n(GBD)', 'Nhóm 15-19 tuổi có tỷ lệ cao hơn 10-14 (5.5% vs 4.1%) [3].\nCho thấy lo âu tăng theo tuổi trong giai đoạn vị thành niên.\nXu hướng tăng 52% toàn cầu từ 1990-2021 [3].\nVùng có chỉ số SDI (Socio-Demographic Index — Chỉ số Nhân khẩu-Xã hội) trung bình (bao gồm VN) chịu ảnh hưởng nặng nhất.'],
    ['Lo âu xã hội\ntoàn cầu\n(2024, PubMed)', 'Meta-\nanalysis', 'Trẻ em', '4.7%\n(chẩn đoán)', 'Lo âu xã hội (social anxiety) là một dạng cụ thể của lo âu.\nỞ trẻ em, tỷ lệ 4.7% — phổ biến nhưng ít được phát hiện.\nLo âu xã hội ảnh hưởng đặc biệt đến khả năng giao tiếp và học tập.\nCần phân biệt với lo âu tổng quát khi sàng lọc tại trường.\nGAD-7 đo lo âu tổng quát, SPIN đo lo âu xã hội (Nakie dùng cả hai).'],
    ['Lo âu xã hội\ntoàn cầu', 'Meta-\nanalysis', 'Thanh\nthiếu niên', '8.3%\n(chẩn đoán)', 'Tỷ lệ lo âu xã hội ở TN (8.3%) gần gấp đôi trẻ em (4.7%).\nTuổi vị thành niên là giai đoạn đỉnh điểm cho lo âu xã hội.\nNakie (2022) báo cáo 37.3% có ám ảnh xã hội (SPIN) — cao hơn nhiều.\nCó thể do SPIN đo triệu chứng rộng hơn tiêu chí chẩn đoán.\nLo âu xã hội thường bị bỏ qua nếu chỉ dùng GAD-7/DASS-21.'],
    ['NSCH Hoa Kỳ\n(2023)', 'N=55,162\nDữ liệu\nquốc gia', '12-17\ntuổi', '16.1%\n(chẩn đoán)', 'Dữ liệu QUỐC GIA Hoa Kỳ — nguồn chính thức liên bang [NSCH, 2020].\nLo âu tăng 61% chỉ trong 7 năm (2016: 10% → 2023: 16.1%).\nNữ 20.1% vs nam 12.3% — chênh lệch giới rõ rệt.\n61% TN gặp khó khăn tiếp cận điều trị dù ở nước phát triển.\nVN (V-NAMHS — Vietnam National Adolescent Mental Health Survey — Khảo sát Quốc gia về Sức khỏe Tâm thần Thanh thiếu niên Việt Nam) 15.6% GẦN TƯƠNG ĐƯƠNG Hoa Kỳ 16.1% — phát hiện đáng chú ý.'],
    ['59 quốc gia\n(2025)\nJ Affect Disord', 'Đa quốc gia\nhọc sinh\nđi học', 'TN\nđi học', '~3.78%\n(ĐNA — Đông Nam Á)', 'Phân tích mới nhất (2025) từ 59 quốc gia [10].\nĐông Nam Á có tỷ lệ lo âu THẤP NHẤT: 3.78%.\nNhưng dữ liệu ĐNA rất hạn chế — ít nghiên cứu được đưa vào.\nTỷ lệ thấp có thể do thiếu sàng lọc hơn là ít bệnh.\nCần nhiều nghiên cứu hơn từ VN và ĐNA để xác minh con số này.'],
]

table = doc.add_table(rows=len(global_data), cols=len(global_data[0]))
table.style = 'Table Grid'
for i, row in enumerate(global_data):
    for j, val in enumerate(row):
        cell = table.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True

doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('Nhận xét: Tỷ lệ lo âu dao động rất rộng từ 3.78% [10] đến 20.5% [1] tùy thuộc vào phương pháp đo lường (chẩn đoán vs sàng lọc), thời điểm (trước vs trong COVID-19), và bối cảnh văn hóa [2][3].')
run.font.size = Pt(10)
run.font.italic = True

doc.add_page_break()

# PART 2: ASEAN
doc.add_heading('PHẦN 2: ĐÔNG NAM Á (ASEAN)', level=1)

doc.add_heading('2.1. Tổng quan khu vực', level=2)
asean_points = [
    'Năm 2021, tỷ lệ rối loạn tâm thần chuẩn hóa theo tuổi ở ASEAN là 11.9% [4].',
    'Lo âu và trầm cảm là 2 rối loạn phổ biến nhất [4].',
    'Nhóm 10-14 tuổi: lo âu ước tính 4.1% [3] — nhưng đây là tỷ lệ rối loạn được chẩn đoán.',
    'Nhóm 10-14 tuổi có gánh nặng bệnh tâm thần CAO NHẤT: 16.3% tổng DALY ở nhóm tuổi này [3].',
    'Tỷ lệ tăng rõ rệt giai đoạn 2019-2021 (COVID-19) [3][4].',
    'Malaysia tăng 12.6%, Indonesia 9.0%, Philippines 7.5% trong giai đoạn 2019-2021 [4].',
]
for point in asean_points:
    doc.add_paragraph(point, style='List Bullet')

doc.add_heading('2.2. Bảng so sánh các nước ASEAN', level=2)

asean_data = [
    ['Quốc gia', 'Nguồn', 'N', 'Nhóm tuổi', 'Lo âu %', 'Công cụ'],
    ['Việt Nam', 'V-NAMHS (2023)', '9,781', '10-17', '15.6%', 'DASS-21'],
    ['Việt Nam', 'Tuy Hòa (2021)', '539', '12-15 (THCS)', '22.7% (RLTT — rối loạn tâm thần chung)', 'SDQ'],
    ['Việt Nam', 'COVID study (2023)', '8,473', 'TN', '41.5% (trong)\n25.4% (sau)', 'DASS-21'],
    ['Việt Nam', 'Dân tộc TT (2024)', '845', 'TN thiểu số', '54.4%', '—'],
    ['Thái Lan', 'GBD 2021', 'Ước tính', '10-24', '~4-5%', 'GBD model'],
    ['Indonesia', 'GBD 2021', 'Ước tính', '10-24', '~4-5%', 'GBD model'],
    ['Philippines', 'GBD 2021', 'Ước tính', '10-24', '~4-5%', 'GBD model'],
    ['Malaysia', 'GBD 2021', 'Ước tính', '10-24', '~5%', 'GBD model'],
    ['Toàn ASEAN', 'GBD 2021', '10 nước', '10-14', '4.1%', 'GBD model'],
    ['Toàn ASEAN', '59 QG (quốc gia) (2025)', 'Đa quốc gia', 'HS đi học', '3.78%', 'Khảo sát'],
]

table2 = doc.add_table(rows=len(asean_data), cols=len(asean_data[0]))
table2.style = 'Table Grid'
for i, row in enumerate(asean_data):
    for j, val in enumerate(row):
        cell = table2.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True
        if 'Việt Nam' in val:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.bold = True

doc.add_page_break()

# PART 3: VIETNAM
doc.add_heading('PHẦN 3: VIỆT NAM — PHÂN TÍCH CHI TIẾT', level=1)

doc.add_heading('3.1. Các nghiên cứu chính về lo âu ở học sinh Việt Nam', level=2)

vn_data = [
    ['Nghiên cứu', 'N', 'Đối tượng', 'Địa điểm', 'Công cụ', 'Lo âu %', 'Phát hiện nổi bật'],
    ['V-NAMHS (2023)', '9,781', 'TN 10-17', 'Toàn quốc', 'DASS-21', '15.6%', 'Đại diện QG đầu tiên.\nTỷ lệ thấp hơn nhiều nước.'],
    ['Tuy Hòa (2021)\nMedPharmRes', '539', 'THCS 12-15', 'Phú Yên', 'SDQ', '22.7%\n(RLTT chung)', 'Nghiên cứu TRỰC TIẾP\nhọc sinh THCS tại VN.'],
    ['COVID VN (2023)', '8,473', 'TN', '6 tỉnh/thành', 'DASS-21', '41.5% (trong)\n25.4% (sau)', 'Giảm 16 điểm %\nsau đại dịch.'],
    ['Dân tộc TT (2024)', '845', 'TN thiểu số', 'Trường nội trú', '—', '54.4%', 'Gấp 3.5 lần V-NAMHS.\nBất bình đẳng nghiêm trọng.'],
    ['TP.HCM (2022)', '384', 'THPT (trung học phổ thông)', 'TP.HCM', 'CES-D', '57.6%\n(trầm cảm)', 'Đô thị lớn nhất.\nÁp lực học tập.'],
    ['Stress học tập (2022)\nCurrent Psychology', '—', 'TN VN', '—', 'Mô hình', '—', 'Resilience = yếu tố\nbảo vệ quan trọng.'],
    ['Hanoi (2023)\nMedicine (LWW)', '5,730', 'HS', 'Hà Nội', '—', '16.2%', 'Tỷ lệ gần V-NAMHS.'],
]

table3 = doc.add_table(rows=len(vn_data), cols=len(vn_data[0]))
table3.style = 'Table Grid'
for i, row in enumerate(vn_data):
    for j, val in enumerate(row):
        cell = table3.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True

doc.add_heading('3.2. Nhận xét về tình hình tại Việt Nam', level=2)

vn_comments = [
    'Tỷ lệ lo âu ở Việt Nam dao động từ 15.6% [5] (V-NAMHS, đại diện quốc gia) đến 54.4% [8] (dân tộc thiểu số). Sự chênh lệch phản ánh bất bình đẳng SKTT nghiêm trọng.',
    'Nghiên cứu TRỰC TIẾP ở học sinh THCS (12-15 tuổi) rất hiếm — chỉ có nghiên cứu Tuy Hòa (2021) với N=539 [6]. Đây là KHOẢNG TRỐNG LỚN cần lấp đầy.',
    'COVID-19 làm tăng lo âu từ ~16% lên 41.5% — gấp 2.5 lần [7]. Sau đại dịch giảm xuống 25.4% nhưng vẫn cao hơn trước [7].',
    'Nhóm dân tộc thiểu số (54.4%) [8] và học sinh trường nội trú cần sự quan tâm ĐẶC BIỆT.',
    'Áp lực học tập và thi cử là yếu tố nguy cơ HÀNG ĐẦU tại VN [9] — phù hợp với Wen (2020) ở nông thôn TQ.',
    'Khả năng phục hồi (resilience) là yếu tố bảo vệ [9] — phù hợp với Qiu (2022) ở TQ.',
    'Thiếu nhân lực SKTT trong trường học — nhân viên y tế trường không được đào tạo sàng lọc SKTT [5] (UNICEF VN).',
    'Thiếu phối hợp giữa trường học và dịch vụ tâm thần địa phương [5].',
]
for c in vn_comments:
    doc.add_paragraph(c, style='List Bullet')

doc.add_page_break()

# PART 4: COMPARISON WITH 11 ORIGINAL PAPERS
doc.add_heading('PHẦN 4: SO SÁNH VỚI 11 BÀI GỐC', level=1)

compare_data = [
    ['Nghiên cứu', 'N', 'Lo âu', 'Phát hiện nổi bật', 'So sánh với VN (V-NAMHS 15.6%)'],
    [
        'Việt Nam\nV-NAMHS\n(2023)',
        '9,781',
        '15.6%\nDASS-21',
        'Khảo sát đại diện quốc gia đầu tiên của VN.\nCỡ mẫu lớn nhất Đông Nam Á.\nTrầm cảm 10%, lo âu 15.6%.\n1/5 thanh thiếu niên VN gặp thách thức SKTT.\nDân tộc thiểu số và nông thôn chênh lệch lớn [5].',
        'Đây là BASELINE quốc gia.\nMọi nghiên cứu VN nên so sánh với con số này.\nTỷ lệ 15.6% phản ánh mẫu đại diện — thấp hơn các nghiên cứu mẫu thuận tiện.'
    ],
    [
        'Việt Nam\nTHCS Tuy Hòa\n(2021)',
        '539',
        '22.7%\n(RLTT chung)\nSDQ',
        'Nghiên cứu DUY NHẤT trực tiếp ở THCS (12-15 tuổi) tại VN.\nSử dụng SDQ — công cụ sàng lọc rối loạn tâm thần chung.\nTỷ lệ 22.7% cao hơn V-NAMHS.\nTại Phú Yên — tỉnh miền Trung [6].',
        'Cao hơn V-NAMHS (22.7% vs 15.6%).\nCó thể do nhóm THCS dễ tổn thương hơn THPT.\nCũng có thể do SDQ đo rộng hơn DASS-21.\nCỡ mẫu nhỏ (N=539) — cần nghiên cứu lớn hơn.'
    ],
    [
        'Trung Quốc\nXu (2021)',
        '373,216',
        '9.89%\nGAD-7',
        'Nghiên cứu LỚN NHẤT toàn cầu về lo âu TN.\nThu thập trong 8 ngày đầu COVID-19 (2/2020).\nNam > nữ (10.11% vs 9.66%) — trái ngược y văn.\nNông thôn cao nhất 12.8% vs thành phố 8.77%.\nSức mạnh thống kê rất cao nhờ cỡ mẫu [Xu, 2021].',
        'VN cao hơn ~1.6 lần (15.6% vs 9.89%).\nNhưng Xu đo tại thời điểm rất sớm COVID-19.\nCông cụ khác nhau (DASS-21 vs GAD-7).\nNếu cùng DASS-21, tỷ lệ TQ có thể cao hơn.\nNông thôn VN có thể tương đồng nông thôn TQ (12.8%).'
    ],
    [
        'Trung Quốc\nChen (2023)',
        '63,205',
        '13.9%\nGAD-7',
        'Nghiên cứu lớn đầu tiên từ miền Tây TQ.\nBắt nạt, giấc ngủ kém, nghiện game là yếu tố nguy cơ mạnh.\nRối loạn giấc ngủ OR=6.99 — yếu tố mạnh nhất.\nNghiện game OR=5.00. Bắt nạt lời nói OR=1.70.\nPhân tích riêng THCS vs THPT [Chen, 2023].',
        'VN gần tương đương (15.6% vs 13.9%).\nCông cụ khác (DASS-21 vs GAD-7) nhưng tỷ lệ gần.\nCác yếu tố nguy cơ (game, giấc ngủ, bắt nạt) có thể tương tự ở VN.\nCần nghiên cứu tương tự tại VN với PHQ-9+GAD-7 để so sánh trực tiếp.'
    ],
    [
        'Trung Quốc\nWen (2020)',
        '900',
        '24.78%\nnặng\nMHT+LPA',
        'Đầu tiên dùng LPA cho lo âu TN nông thôn TQ.\n3 nhóm lo âu: thấp, trung bình, nặng (24.78%).\nNAM > nữ — trái ngược y văn, phù hợp Saikia.\nHỗ trợ SKTT tại trường là yếu tố bảo vệ.\nÁp lực học tập cực kỳ lớn ở nông thôn [Wen, 2020].',
        'Nông thôn kém phát triển TQ ~ nông thôn VN.\nTrẻ bị bỏ lại (left-behind) ở TQ ~ trẻ có bố mẹ đi làm xa ở VN.\nHỗ trợ SKTT tại trường — VN hiện thiếu nhân lực [5].\nCần nghiên cứu LPA tương tự ở nông thôn VN.'
    ],
    [
        'Ấn Độ\nSaikia (2023)',
        '360',
        '24.4%\nDASS-21',
        'Đầu tiên từ Đông Bắc Ấn Độ.\nNAM > nữ (30% vs 18.9%) — trái ngược y văn.\nCha mẹ uống rượu: trầm cảm 38.3% vs 16.5%.\nChơi game: trầm cảm 27.9% vs 10.8%.\nBị lưu ban: trầm cảm 58.6% [Saikia, 2023].',
        'VN thấp hơn 1.6 lần (15.6% vs 24.4%) — cùng DASS-21!\nĐây là so sánh TRỰC TIẾP nhất (cùng công cụ).\nCó thể do VN có hệ thống giáo dục ổn định hơn.\nYếu tố rượu cha mẹ cũng cần khảo sát ở VN.'
    ],
    [
        'Ấn Độ\nBhardwaj\n(2020)',
        '288',
        '73.3%\nDASS-21',
        'Tỷ lệ CAO NHẤT trong tất cả nghiên cứu.\nLo âu nặng + cực nặng: 46.8%.\nTrường công lập, học sinh kinh tế thấp.\nChỉ thống kê mô tả — không phân tích yếu tố.\nTạp chí địa phương, chất lượng phương pháp thấp [Bhardwaj, 2020].',
        'VN thấp hơn 4.7 lần (15.6% vs 73.3%) — cùng DASS-21!\nChênh lệch cực lớn dù cùng công cụ → bối cảnh rất khác.\nTrường công VN có thể tương đồng trường công Ấn Độ.\nCần so sánh trường công vs trường tư tại VN.'
    ],
    [
        'Ethiopia\nNakie (2022)',
        '849',
        '66.7%\nDASS-21',
        'Đầu tiên ở châu Phi đánh giá cả 3 tình trạng.\nKhat AOR=5.6, hút thuốc AOR=4.8 — yếu tố mạnh.\nBối cảnh xung đột Tigray 2020-2022.\n46.5% đã dùng rượu ở lứa tuổi đi học.\nPhạm vi tuổi rộng 15-25 [Nakie, 2022].',
        'VN thấp hơn 4.3 lần (15.6% vs 66.7%).\nBối cảnh xung đột Ethiopia không có ở VN.\nNhưng sử dụng chất (rượu, thuốc) ở TN VN cần khảo sát.\nCùng DASS-21 — chênh lệch phản ánh bối cảnh xã hội.'
    ],
    [
        'Ả Rập Saudi\nAlharbi (2019)',
        '1,245',
        '63.5%\nGAD-7',
        'Tỷ lệ rất cao: 74% trầm cảm, 63.5% lo âu.\nĐầu tiên ở Saudi dùng PHQ-9/GAD-7.\nNữ > nam (P<0.001). Buraidah chiếm ưu thế.\nNgưỡng PHQ-9 ≥5 bao gồm cả nhẹ → phóng đại.\nĐược trích dẫn rộng rãi khu vực Trung Đông [Alharbi, 2019].',
        'VN thấp hơn 4 lần (15.6% vs 63.5%).\nNhưng công cụ khác (DASS-21 vs GAD-7).\nNgưỡng cắt khác nhau → so sánh cần thận trọng.\nÁp lực văn hóa và kỳ vọng gia đình Saudi khác VN.'
    ],
    [
        'Hoa Kỳ\nNSCH (2020)',
        '55,162',
        '16.1%\nChẩn đoán',
        'Dữ liệu quốc gia Hoa Kỳ, uy tín rất cao.\nLo âu tăng 61% từ 2016-2023 (10% → 16.1%).\nNữ 20.1% vs nam 12.3%.\n61% gặp khó khăn tiếp cận điều trị.\nDựa trên chẩn đoán y tế — khác sàng lọc [NSCH, 2020].',
        'VN gần TƯƠNG ĐƯƠNG (15.6% vs 16.1%)!\nĐây là so sánh đáng chú ý nhất.\nNhưng NSCH dùng chẩn đoán, VN dùng sàng lọc.\nNếu VN dùng chẩn đoán, tỷ lệ có thể THẤP HƠN 15.6%.\nXu hướng tăng ở Mỹ cần theo dõi ở VN.'
    ],
    [
        'Toàn cầu\nRacine (2021)',
        '>80,000',
        '20.5%\nMeta-analysis',
        'Meta-analysis QUAN TRỌNG NHẤT về COVID-19.\n1/4 TN có trầm cảm, 1/5 có lo âu.\nGấp đôi so với trước đại dịch.\nNữ và TN lớn tuổi bị ảnh hưởng nhiều hơn.\nĐược trích dẫn >2,000 lần [1].',
        'VN sau COVID (25.4%) > trung bình toàn cầu (20.5%).\nVN trong COVID (41.5%) > trung bình gấp 2 lần.\nSau đại dịch VN giảm nhưng vẫn cao hơn baseline.\nCần theo dõi dài hạn tác động hậu COVID ở VN.'
    ],
]

table4 = doc.add_table(rows=len(compare_data), cols=len(compare_data[0]))
table4.style = 'Table Grid'
for i, row in enumerate(compare_data):
    for j, val in enumerate(row):
        cell = table4.cell(i, j)
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(8)
        run.font.name = 'Times New Roman'
        if i == 0:
            run.bold = True
        if 'Việt Nam' in val:
            run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
            run.bold = True

doc.add_page_break()

# PART 5: GAPS AND RECOMMENDATIONS
doc.add_heading('PHẦN 5: KHOẢNG TRỐNG VÀ ĐỀ XUẤT NGHIÊN CỨU', level=1)

doc.add_heading('5.1. Khoảng trống tại Việt Nam', level=2)
gaps = [
    'Rất THIẾU nghiên cứu trực tiếp ở nhóm THCS (12-15 tuổi) — chỉ có 1 bài (Tuy Hòa, N=539) [6].',
    'Chưa có nghiên cứu dọc (longitudinal) theo dõi lo âu theo thời gian ở VN — tất cả đều cắt ngang [5][6][7].',
    'Chưa có nghiên cứu can thiệp (intervention) đánh giá hiệu quả chương trình SKTT tại trường ở VN [5].',
    'Thiếu dữ liệu từ vùng nông thôn, miền núi, dân tộc thiểu số — V-NAMHS cho thấy chênh lệch lớn [5][8].',
    'Chưa xác thực đầy đủ công cụ sàng lọc (GAD-7, PHQ-9) cho nhóm THCS 12-15 tuổi tại VN — Nam Phi đã làm cho GAD-7/PHQ-9 [11].',
    'Thiếu nghiên cứu vai trò mạng xã hội, game trực tuyến đối với lo âu ở VN — Chen (2023) đã chứng minh IGD OR=5.00 ở TQ.',
]
for g in gaps:
    doc.add_paragraph(g, style='List Bullet')

doc.add_heading('5.2. Đề xuất nghiên cứu ưu tiên', level=2)
recs = [
    ('Ưu tiên 1:', 'Nghiên cứu cắt ngang quy mô lớn (N>1,000) về lo âu ở học sinh THCS tại nhiều tỉnh/thành VN, sử dụng DASS-21 hoặc GAD-7 (phiên bản tiếng Việt đã xác thực).'),
    ('Ưu tiên 2:', 'So sánh thành thị vs nông thôn, Kinh vs dân tộc thiểu số — V-NAMHS cho thấy bất bình đẳng lớn cần khám phá sâu.'),
    ('Ưu tiên 3:', 'Nghiên cứu can thiệp pilot — chương trình SKTT tại trường (school-based) với đánh giá hiệu quả trước-sau.'),
    ('Ưu tiên 4:', 'Nghiên cứu dọc theo dõi nhóm THCS qua 2-3 năm để xác lập xu hướng và yếu tố nhân quả.'),
    ('Ưu tiên 5:', 'Xác thực công cụ sàng lọc (GAD-7, PHQ-9A, DASS-21) cho nhóm THCS 12-15 tuổi tại VN — tương tự nghiên cứu Nam Phi (2022).'),
]
for label, content in recs:
    p = doc.add_paragraph()
    run = p.add_run(label + ' ')
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x00, 0x66, 0x00)
    run = p.add_run(content)
    run.font.size = Pt(11)

doc.add_paragraph()
doc.add_heading('5.3. Nguồn tham khảo chính', level=2)
refs = [
    '1. Racine N et al. (2021). JAMA Pediatrics, 175(11):1142-1150. Meta-analysis COVID-19.',
    '2. Polanczyk GV et al. (2015). J Child Psychol Psychiatry, 56:345-365. Meta-analysis toàn cầu.',
    '3. GBD 2021 Collaborators. Translational Psychiatry (2025). Gánh nặng bệnh toàn cầu.',
    '4. GBD 2021 ASEAN. Lancet Public Health (2025). Gánh nặng bệnh ASEAN.',
    '5. V-NAMHS (2023). UNICEF/QCMHR. Khảo sát quốc gia VN.',
    '6. Tuy Hòa THCS (2021). MedPharmRes, 7(2):32. Nghiên cứu THCS VN.',
    '7. COVID VN (2023). Am J Psychiatric Rehabilitation. DASS-21 6 tỉnh VN.',
    '8. Dân tộc thiểu số VN (2024). Mental Health & Prevention. Elsevier.',
    '9. Anderson et al. (2025). J Child Adolesc Psychiatric Nursing. Review yếu tố nguy cơ.',
    '10. 59 quốc gia (2025). J Affect Disord. Phân tích đa quốc gia lo âu học sinh.',
    '11. Lancet Regional Health SEA (2025). Review Nam Á + ĐNA.',
]
for ref in refs:
    p = doc.add_paragraph(ref)
    p.runs[0].font.size = Pt(9)

out = os.path.join(BASE, 'DocFiles', 'BAO_CAO_LO_AU_THCS_VN_THE_GIOI.docx')
doc.save(out)
print(f'DONE: {out}')
