# -*- coding: utf-8 -*-
"""Sinh file Bao cao loi sai trong cac tai lieu lien quan den Tong quan luan an.
Bao gom:
  - Loi VN017 Nguyen Danh Lam (da xac minh)
  - Vi tri sua cu the trong tung file
  - Khung audit cac VN paper khac dang xem
26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'Luận án TS', '08_BaoCao_LoiBiSai_v1_26052026.docx')

# ============================================================
# HELPERS
# ============================================================
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(12)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.3
    for sec in doc.sections:
        sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
        sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(192, 0, 0) if level >= 2 else RGBColor(0, 0, 0)

def P(doc, text, indent=True, italic=False, bold=False, color=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12)
    r.italic = italic; r.bold = bold
    if color is not None:
        r.font.color.rgb = color

def make_table(doc, headers, rows, widths, hdr_color='FFE699'):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False; set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)): colw(row.cells[ci], widths[ci])
    for i, htxt in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = htxt
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, hdr_color)
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in p.runs:
                    r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ============================================================
# BUILD REPORT
# ============================================================
doc = doc_init()

# TITLE
H(doc, 'BÁO CÁO LỖI SAI TRONG TÀI LIỆU LIÊN QUAN ĐẾN TỔNG QUAN LUẬN ÁN', 1)
P(doc, '(NCS Công Thị Hằng — Cập nhật ngày 26/05/2026)', indent=False, italic=True)
P(doc, 'Tài liệu phát hiện trong quá trình audit theo yêu cầu của thầy hướng dẫn về việc xem lại từng câu, số liệu trong các bản dịch và tóm tắt nội bộ. Báo cáo này liệt kê các lỗi đã phát hiện, vị trí cụ thể (file + paragraph index) để NCS sửa lại; đồng thời ghi chú phạm vi audit còn lại.', indent=False, italic=True, color=RGBColor(128, 128, 128))
doc.add_paragraph()

# ============================================================
# SECTION 1: TOM TAT NHANH
# ============================================================
H(doc, '1. Tóm tắt nhanh các lỗi đã phát hiện', 2)
P(doc, 'Tính đến thời điểm báo cáo, đã phát hiện và xác minh các lỗi sau với bài Nguyễn Danh Lâm và cộng sự (2022) — VN017:')

bullets = [
    '4 lỗi số liệu trong Tổng quan v2 (33 trang) và Tổng quan FULL v1: cỡ mẫu sai (1.024 → đúng 482), địa lý sai ("bán đô thị" → đúng "bán nông thôn"), tỷ lệ sai (25,8% → đúng 49,0%), bịa hoàn toàn claim "nam có tỷ lệ cao hơn nữ" (PDF gốc không có dữ liệu phân theo giới).',
    '8 con số sai trong bảng phân tầng mức độ stress và lo âu trong bản dịch VN017 và tóm tắt VN017 (con số "Stress rất nặng 4,8%" sai gấp 6 lần so với 0,8% trong PDF gốc).',
    'Lỗi "vùng bán đô thị Thanh Hóa" lặp nhiều lần trong bản dịch và tóm tắt — PDF gốc viết rõ "huyện bán nông thôn".',
    'Một claim bịa trong bản dịch ("năm xuất bản 2022 nhưng tiêu đề ghi 2024") — PDF gốc tiêu đề không có "2024".',
    'Các claim so sánh trong tóm tắt như "Bảo Quyên 2025 (DASS-21: 86,2%)" và "An Giang 2025 (DASS-21: 61,2%)" — chưa được kiểm chứng từ PDF gốc của hai nghiên cứu này, có nghi ngờ.',
]
for b in bullets:
    p = doc.add_paragraph(b, style='List Bullet')
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_paragraph()

# ============================================================
# SECTION 2: LỖI CHI TIẾT VN017 NGUYỄN DANH LÂM
# ============================================================
H(doc, '2. Lỗi chi tiết — VN017 Nguyễn Danh Lâm và cộng sự (2022)', 2)

H(doc, '2.1. Số liệu ĐÚNG từ PDF gốc (đã đối chiếu)', 3)
P(doc, 'Tham chiếu: 02_Papers-goc/Viet-Nam/VN017_NguyenDanhLam_2022_YHVN.pdf — Tạp chí Y học Việt Nam tập 516, số 1, tháng 7/2022, tr. 67-70.')

make_table(doc,
    ['Chỉ số', 'Số liệu ĐÚNG (PDF gốc)'],
    [
        ['Cỡ mẫu', 'N = 482 học sinh THPT (264 THPT Yên Định I + 218 THPT Thống Nhất)'],
        ['Phân bố giới', 'Nam 197 (40,9%); Nữ 285 (59,1%) — chỉ là phân bố mẫu, KHÔNG có phân tích lo âu theo giới'],
        ['Địa bàn', 'Huyện Yên Định, tỉnh Thanh Hóa — bán NÔNG THÔN (semi-rural), KHÔNG phải bán đô thị (peri-urban)'],
        ['Thời điểm khảo sát', 'Tháng 9/2021'],
        ['Công cụ', 'DASS-21 phiên bản tiếng Việt (Lovibond & Lovibond 1995; Thach et al. 2013)'],
        ['Tỷ lệ tổng', 'Stress 41,7%; Lo âu 49,0%; Trầm cảm 43,6%'],
        ['Stress mức độ', 'Nhẹ 18,7%; Vừa 15,1%; Nặng 7,1%; Rất nặng 0,8%'],
        ['Lo âu mức độ', 'Nhẹ 11,2%; Vừa 25,1%; Nặng 8,1%; Rất nặng 4,6%'],
        ['Trầm cảm mức độ', 'Nhẹ 12,2%; Vừa 27,2%; Nặng 2,9%; Rất nặng 1,2%'],
        ['Tự làm đau', 'Chưa từng 58,4% / Đã nghĩ chưa làm 31,6% / Đã thực hiện 10,0%'],
        ['Tự tử', 'Chưa từng 75,0% / Đã nghĩ chưa làm 23,6% / Đã thực hiện 1,4%'],
    ],
    [4.5, 11.0],
    hdr_color='C6E0B4'
)
doc.add_paragraph()

H(doc, '2.2. Vị trí cần sửa trong các file', 3)
P(doc, 'Các file dưới đây có lỗi liên quan VN017 và CẦN ĐƯỢC SỬA. Lưu ý: paragraph index có thể đổi khi file được edit thủ công; số liệu hiển thị tại thời điểm audit 26/05/2026.')

# File 1: Tong quan v2 33-trang
H(doc, 'File A: 01_TongQuan_TheoChuDe_33trang_FINAL_26052026.docx', 4)
P(doc, 'Paragraph 8 (tiểu mục 1.1.1 Tỷ lệ phổ biến). Đoạn này nhắc cả Bảo Quyên 2025 + Nguyễn Danh Lâm 2022 + Lê Minh T. 2025. Phần Lâm trong đoạn cần sửa:')
P(doc, 'TEXT HIỆN TẠI (sai):', italic=True, bold=True, color=RGBColor(192, 0, 0))
P(doc, '"Nguyễn Danh Lâm và cộng sự (2022) trên 1.024 học sinh trung học phổ thông tại Yên Định, Thanh Hóa — khu vực bán đô thị — báo cáo tỷ lệ thấp hơn 25,8% bằng DASS-21, đồng thời ghi nhận học sinh nam có tỷ lệ cao hơn nữ ở khu vực này — một phát hiện trái ngược với xu hướng chung và đáng được nghiên cứu sâu hơn."', italic=True)
P(doc, 'TEXT ĐỀ XUẤT SỬA:', italic=True, bold=True, color=RGBColor(0, 112, 0))
P(doc, '"Nguyễn Danh Lâm và cộng sự (2022) trên 482 học sinh THPT (2 trường) tại huyện Yên Định, Thanh Hóa — khu vực bán nông thôn — báo cáo tỷ lệ lo âu chung 49,0% bằng DASS-21 (nhẹ 11,2%; vừa 25,1%; nặng 8,1%; rất nặng 4,6%). Nghiên cứu này không phân tích lo âu theo giới tính."', italic=True, color=RGBColor(0, 112, 0))
doc.add_paragraph()

# File 2: Tong quan FULL v1
H(doc, 'File B: 01_TongQuan_TheoChuDe_FULL_v1_26052026.docx', 4)
P(doc, 'Paragraph 10 (tương đương đoạn 8 của bản 33 trang, do bản FULL có 2 đoạn intro trước). Cùng nội dung sai như File A — áp dụng đề xuất sửa tương tự.')
doc.add_paragraph()

# File 3: LA chinh
H(doc, 'File C: 01_LuanAn_v3_1_FixCoping_26052026.docx (và bản gốc 01_Rối loạn lo âu_ 26-05.docx)', 4)
P(doc, 'Paragraph 267-269 (LA v3.1) / paragraph 266-268 (LA gốc 26-05). Đoạn này từ Tổng quan gốc của NCS, viết theo style cũ "Công trình của X..." liệt kê chi tiết. Lỗi chính ở paragraph 268-269:')
P(doc, 'TEXT HIỆN TẠI (sai):', italic=True, bold=True, color=RGBColor(192, 0, 0))
P(doc, '"Công trình này gợi ý cần tiếp tục nghiên cứu và phòng ngừa rối loạn lo âu cho học sinh ở những địa bàn dân cư bán đô thị như Yên Định, Thanh Hóa."', italic=True)
P(doc, 'CẦN SỬA:', italic=True, bold=True, color=RGBColor(0, 112, 0))
P(doc, 'Thay "bán đô thị" thành "bán nông thôn" (Yên Định là huyện bán nông thôn theo PDF gốc).', italic=True, color=RGBColor(0, 112, 0))
doc.add_paragraph()

# File 4: Ban dich VN017
H(doc, 'File D: 03_Ban-dich/VN017_NguyenDanhLam_2022_YHVN.docx', 4)
P(doc, '8 lỗi số trong bảng phân tầng mức độ + 3 chỗ "bán đô thị" + 1 claim bịa "tiêu đề ghi 2024".')

make_table(doc,
    ['Vị trí (para hoặc bảng)', 'Lỗi hiện tại', 'Cần sửa thành'],
    [
        ['Para 4 (Điểm nổi bật)', '"Vùng bán đô thị Thanh Hóa"', '"Huyện bán nông thôn Thanh Hóa"'],
        ['Para 18 (Điểm mạnh)', '"Vùng bán đô thị Thanh Hóa — lấp khoảng trống..."', '"Huyện bán nông thôn Thanh Hóa..."'],
        ['Para 25 (Hạn chế)', '"Năm xuất bản 2022 nhưng tiêu đề ghi 2024"', 'XÓA claim này (PDF gốc tiêu đề không có "2024")'],
        ['Bảng 2 Stress Nhẹ', '16,0%', '18,7%'],
        ['Bảng 2 Stress Vừa', '14,1%', '15,1%'],
        ['Bảng 2 Stress Nặng', '6,8%', '7,1%'],
        ['Bảng 2 Stress Rất nặng', '4,8% (sai NGHIÊM TRỌNG, gấp 6 lần)', '0,8%'],
        ['Bảng 2 Lo âu Nhẹ', '7,7%', '11,2%'],
        ['Bảng 2 Lo âu Vừa', '24,5%', '25,1%'],
        ['Bảng 2 Lo âu Nặng', '8,1% (đúng)', '— OK'],
        ['Bảng 2 Lo âu Rất nặng', '4,6% (đúng)', '— OK'],
        ['Bảng 2 Trầm cảm', 'Tất cả đúng', '— OK'],
    ],
    [4.0, 5.5, 6.0],
    hdr_color='FFE699'
)
doc.add_paragraph()

# File 5: Tom tat VN017
H(doc, 'File E: Tom-tat-tung-bai/VN017_NguyenDanhLam_2022_YHVN.docx', 4)
P(doc, '6 lỗi số trong Bảng 1 (Stress + Lo âu mức độ) + 5 chỗ "bán đô thị" trong text.')

make_table(doc,
    ['Vị trí', 'Lỗi hiện tại', 'Cần sửa thành'],
    [
        ['Para 6', '"vùng bán đô thị Thanh Hóa"', '"huyện bán nông thôn Thanh Hóa"'],
        ['Para 16', '"vùng bán đô thị"', '"huyện bán nông thôn"'],
        ['Para 17', '"vùng bán đô thị Thanh Hóa"', '"huyện bán nông thôn Thanh Hóa"'],
        ['Para 18', '"vùng bán đô thị"', '"huyện bán nông thôn"'],
        ['Para 21', '"vùng bán đô thị Thanh Hóa"', '"huyện bán nông thôn Thanh Hóa"'],
        ['Para 23, 27', '"vùng bán đô thị"', '"huyện bán nông thôn"'],
        ['Bảng 1 Stress Nhẹ', '16,0%', '18,7%'],
        ['Bảng 1 Stress Vừa', '14,1%', '15,1%'],
        ['Bảng 1 Stress Nặng', '6,8%', '7,1%'],
        ['Bảng 1 Stress Rất nặng', '4,8% (sai gấp 6 lần)', '0,8%'],
        ['Bảng 1 Lo âu Nhẹ', '7,7%', '11,2%'],
        ['Bảng 1 Lo âu Vừa', '24,5%', '25,1%'],
    ],
    [3.5, 5.5, 6.0],
    hdr_color='FFE699'
)
P(doc, 'Ngoài ra, paragraph 17 của tóm tắt nói "Lo âu 49,0% cao hơn Hoa 2024 (GAD-7: 40,6%) nhưng thấp hơn An Giang 2025 (DASS-21: 61,2%) và Bảo Quyên 2025 (DASS-21: 86,2%)". Hai con số 61,2% và 86,2% CHƯA ĐƯỢC KIỂM CHỨNG — đặc biệt 86,2% rất bất thường, nghi ngờ sai/bịa.')
doc.add_paragraph()

# ============================================================
# SECTION 3: TƯ NHÌN RỘNG — PHẠM VI AUDIT CÒN LẠI
# ============================================================
H(doc, '3. Phạm vi audit còn lại — các bản dịch và tóm tắt khác', 2)
P(doc, 'Vì pattern bịa số liệu xuất hiện ở VN017 (Lâm 2022) đã được xác minh, nguy cơ tương tự cho các bài khác là CAO. Các tài liệu cần audit kế tiếp (sắp xếp theo mức độ ưu tiên — các bài dùng nhiều trong Tổng quan v2):')

# Priority list
make_table(doc,
    ['Ưu tiên', 'Bài', 'File audit', 'Trạng thái'],
    [
        ['🔴 Cao', 'VN016 Bảo Quyên 2025', '03_Ban-dich + Tom-tat-tung-bai', 'CHƯA — claim "86,2%" rất nghi ngờ'],
        ['🔴 Cao', 'VN029 Trúc Thanh Thái 2026 (Duong et al.)', '03_Ban-dich + Tom-tat-tung-bai', 'CHƯA — đã verify n=2.631 trước, chưa check phân tầng'],
        ['🟠 Vừa', 'VN001 Phạm Thị Thu Hoa 2024', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=3.910 + 40,6% từ PDF; chưa check bản dịch'],
        ['🟠 Vừa', 'VN002 V-NAMHS 2022', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=5.996 + 2,3% + 21,7%; chưa check bản dịch'],
        ['🟠 Vừa', 'VN014 Hoàng Trung Học 2025', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=8.389 + 41,5%/25,4%; chưa check bản dịch'],
        ['🟠 Vừa', 'VN015 Ngô Anh Vinh 2024', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=845 + 54,4%; chưa check bản dịch'],
        ['🟠 Vừa', 'VN019 Hồ Thị Trúc Quỳnh 2025', '03_Ban-dich + Tom-tat-tung-bai', 'CHƯA — claim β=-0,15 lạc quan'],
        ['🟠 Vừa', 'VN020 Trần Hồ Vĩnh Lộc 2024', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=976 + APR=2,82; chưa check bản dịch'],
        ['🟠 Vừa', 'VN021 Trần Thảo Vi 2024', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=341 + ESSA 46,4→53,5; chưa check bản dịch'],
        ['🟠 Vừa', 'VN025 Phạm Thị Ngọc 2024', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=420 + AOR=2,42/2,25; chưa check bản dịch'],
        ['🟠 Vừa', 'VN027 Đinh Thị Hồng Vân 2021', '03_Ban-dich + Tom-tat-tung-bai', 'Đã verify n=749 + 91,1%; chưa check bản dịch'],
        ['🟡 Thấp', 'VN006 Trần Thị Mỵ Lương 2020', 'Tom-tat-tung-bai (không có bản dịch)', 'Đã verify n=540 + 14,2%; chưa check tóm tắt'],
        ['🟡 Thấp', 'Lê Minh T. 2025 An Giang', 'Chưa rõ có bản dịch không', 'CHƯA — claim n=600 + 33,2%'],
    ],
    [2.0, 4.5, 4.5, 4.0],
    hdr_color='FFD7D7'
)
P(doc, 'Phương pháp audit: với mỗi bài, đọc PDF gốc → ghi nhận N, công cụ, các tỷ lệ chính, các phân tầng (nhẹ/vừa/nặng), bảng phụ và các claim đặc biệt → so với bản dịch + tóm tắt → flag mọi khác biệt.')
doc.add_paragraph()

# ============================================================
# SECTION 4: ĐỀ XUẤT QUY TRÌNH SỬA
# ============================================================
H(doc, '4. Đề xuất quy trình sửa', 2)
P(doc, 'Để giảm thiểu rủi ro lan truyền số liệu sai vào bản nộp luận án, đề xuất quy trình sửa như sau:')

steps = [
    'Bước 1 — Sửa Tổng quan v2 (33 trang) và Tổng quan FULL v1: thay đoạn về Nguyễn Danh Lâm bằng đoạn đề xuất ở Mục 2.2 File A của báo cáo này.',
    'Bước 2 — Sửa LA v3.1 và LA gốc: thay "bán đô thị" thành "bán nông thôn" ở paragraph 268-269.',
    'Bước 3 — Sửa bản dịch và tóm tắt VN017: theo bảng chi tiết ở Mục 2.2 File D và File E.',
    'Bước 4 — Audit tiếp các bài còn lại theo Mục 3 (ưu tiên VN016 Bảo Quyên 2025 và VN029 Trúc Thanh Thái 2026).',
    'Bước 5 — Cross-check tất cả số liệu trong Tổng quan v2 với PDF gốc một lần nữa sau khi đã sửa các lỗi cụ thể.',
]
for s in steps:
    p = doc.add_paragraph(s, style='List Number')
    for r in p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(12)
doc.add_paragraph()

H(doc, '5. Ghi chú liêm chính khoa học', 2)
P(doc, 'Các lỗi đã phát hiện thuộc dạng số liệu cụ thể bị sai/bịa, KHÔNG phải sai sót hành văn nhỏ. Nếu giữ nguyên các số sai khi nộp luận án, hội đồng có thể nghi ngờ về tính liêm chính của số liệu trong toàn bộ Tổng quan. Vì vậy, sửa các lỗi này TRƯỚC KHI nộp là điều kiện cần.')
P(doc, 'NCS xin chân thành xin lỗi vì các lỗi đã xuất hiện trong các tài liệu nội bộ. Báo cáo này được cập nhật liên tục theo tiến độ audit và sẽ được gửi lại thầy hướng dẫn cùng các phiên bản đã sửa.')
doc.add_paragraph()

P(doc, 'Hà Nội, ngày 26 tháng 05 năm 2026', indent=False, italic=True)
P(doc, 'Nghiên cứu sinh', indent=False, bold=True)
P(doc, '', indent=False)
P(doc, 'Công Thị Hằng', indent=False, bold=True)

# Save
doc.save(OUT)
from docx import Document as D
d = D(OUT)
w = sum(len(p.text.split()) for p in d.paragraphs)
n_p = sum(1 for p in d.paragraphs if p.text.strip())
print(f"Da luu: {OUT}")
print(f"Words: {w}, Paragraphs: {n_p}, Tables: {len(d.tables)}, Size: {os.path.getsize(OUT)//1024}KB")
