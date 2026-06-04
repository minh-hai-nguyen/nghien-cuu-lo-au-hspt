"""
STEP 3: Canonical 3 bài QT ưu tiên (QT018-020) + Jenkins summary.
- QT018 Salari 2024 SAD Global SR/MA
- QT019 Shibuya 2025 School MH Literacy 3 Asian
- QT020 Liu 2025 CBT Delivery GAD NMA (NGƯỜI LỚN — cảnh báo scope)
"""
import sys, io, os, json, shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
QT_FOLDER_KHAC = BASE / '02_Papers-goc' / 'The-gioi_Khac'
TT_DIR = BASE / 'Tom-tat-tung-bai'

PAPERS = [
    {
        'id': 'QT018',
        'descriptor': 'Salari_SAD_Prevalence_Global_SR_MA_2024',
        'pdf_src': 's10935-024-00789-9.pdf',
        'pdf_dst': 'QT018_Salari_SAD_Prevalence_Global_SR_MA_2024.pdf',
        'folder': 'The-gioi_Khac',
        'title_vn': 'Tỷ lệ mắc rối loạn lo âu xã hội toàn cầu ở trẻ em, vị thành niên và thanh niên: Tổng quan hệ thống và phân tích tổng hợp',
        'title_en': 'Global Prevalence of Social Anxiety Disorder in Children, Adolescents and Youth: A Systematic Review and Meta-analysis',
        'authors': 'Salari N, Heidarian P, Hassanabadi M, Babajani F, Abdoli N, Aminian M, Mohammadi M',
        'affiliation': 'Nhóm Iran + Ả-rập Saudi (hệ thống đại học quốc gia)',
        'journal': 'Journal of Prevention (2024) 45:795-813 · Springer',
        'year': 2024,
        'sample': '38 nghiên cứu nguyên bản gộp trong MA (tổng n không báo trong abstract)',
        'location': 'Toàn cầu (6 databases: PubMed, Scopus, Web of Science, Embase, ScienceDirect, Google Scholar)',
        'tool': 'Random-effects MA; Comprehensive Meta-Analysis v2.0; I² heterogeneity',
        'design': 'Systematic review + meta-analysis (random-effects)',
        'outcomes': [
            'Tỷ lệ SAD toàn cầu TRẺ EM: 4,7 %',
            'Tỷ lệ SAD toàn cầu VỊ THÀNH NIÊN: 8,3 %',
            '(Tỷ lệ ở thanh niên — số cụ thể cần verify chương kết quả)',
            '38 nghiên cứu, random-effects — có tính heterogeneity qua I²',
        ],
        'strengths': [
            'Search 6 databases lớn (toàn diện)',
            'Tách theo 3 nhóm tuổi (children / adolescents / youth) — đặc biệt hữu ích cho đề tài HS THCS/THPT',
            'J Prevention — tạp chí peer-reviewed có IF',
            'Methodology MA chuẩn (random-effects, Comprehensive MA v2.0)',
        ],
        'limitations': [
            'Nhóm tác giả Iran — cần check conflict of interest / selection bias ngôn ngữ',
            'Số liệu SAD children 4,7 % thấp hơn nhiều NC cá thể (ví dụ Jefferies 2020 VN 30,7 %) — có thể do heterogeneity cao hoặc hợp nhất phương pháp',
            'Chưa đọc chi tiết forest plot — không rõ heterogeneity I²',
            'Abstract không nêu rõ số cỡ mẫu tổng (chỉ số NC)',
        ],
        'critique_short': 'MA global về SAD ở trẻ em/VTN — QUAN TRỌNG để có baseline toàn cầu so sánh VN. 4,7 % children + 8,3 % adolescents là ước tính MA "chặt" nhưng thấp hơn các NC cắt ngang VN (Jefferies 2020: VN 30,7 %). Chênh lệch lớn đáng chú ý — có thể do: (1) SIAS-17 đo lo âu tương tác khác với chẩn đoán SAD đầy đủ; (2) heterogeneity giữa các NC trong MA.',
        'future': 'Subgroup analysis theo khu vực (LMIC vs HIC, Đông Á vs Nam Á vs Đông Nam Á). Đối chiếu với Jefferies 2020 (7 nước). Verify số cụ thể "youth" từ chương kết quả.',
        'rating': 4,
        'notes_for_rag': 'Global SAD: 4,7 % children, 8,3 % VTN (Salari 2024 MA 38 NC). Dùng để compare baseline với VN 30,7 %.',
    },
    {
        'id': 'QT019',
        'descriptor': 'Shibuya_SchoolMH_Literacy_3Asian_2025',
        'pdf_src': 's41182-025-00697-6.pdf',
        'pdf_dst': 'QT019_Shibuya_SchoolMH_Literacy_3Asian_2025.pdf',
        'folder': 'The-gioi_Khac',
        'title_vn': 'Nghiên cứu so sánh về hiểu biết sức khỏe tâm thần trường học ở ba nước châu Á',
        'title_en': 'Comparative study on school-based mental health literacy in three Asian countries',
        'authors': 'Shibuya F, Usami M, Santillan MD, Warnaini C, Gregorio E, Satake N, Estrada CA, Gunawan G, Balderrama N, De Leon JF, Ancheta JF, Kadriyan H, Garcia F, Kobayashi J',
        'affiliation': 'Multi-site: Nhật Bản (Ryukyu), Philippines, Indonesia',
        'journal': 'Tropical Medicine and Health (2025) 53:86 · Springer Open',
        'year': 2025,
        'sample': 'Phân tích tài liệu chính sách + chương trình học (lớp 1-12) từ 3 quốc gia Philippines, Indonesia, Nhật Bản, giai đoạn 2000-2023',
        'location': 'Philippines + Indonesia + Japan',
        'tool': 'Deductive content analysis · Policy triangle framework (Walt & Gilson 1994) · Review points (Margaretha 2023)',
        'design': 'Nghiên cứu tài liệu (documentary analysis) — so sánh chính sách + curriculum 3 nước',
        'outcomes': [
            'Ba quốc gia có cách tiếp cận school-based MH literacy KHÁC NHAU về chính sách và curriculum',
            'Phân tích theo policy triangle (context, content, process, actors)',
            'Curriculum lớp 1-12 được map theo các review points',
            '(Kết quả chi tiết cần verify chương Results — abstract chỉ nêu phương pháp)',
        ],
        'strengths': [
            'Nghiên cứu SO SÁNH 3 nước châu Á — hiếm có',
            'Khung policy triangle được sử dụng hệ thống',
            'Multi-stakeholder authors từ cả 3 nước → perspective balance',
            'Scope rộng: 2000-2023, lớp 1-12 — cover toàn bộ học đường',
            'Có tính ngoại suy cho bối cảnh VN (tương tự SEA)',
        ],
        'limitations': [
            'Chỉ phân tích TÀI LIỆU, không đo lường thực tế triển khai',
            'VN không nằm trong 3 nước → cần ngoại suy thận trọng',
            'Nghiên cứu chính sách/curriculum — không đo outcome SKTT HS',
            'Tropical Medicine and Health — không phải tạp chí chuyên MH top',
        ],
        'critique_short': 'Nghiên cứu về CHÍNH SÁCH school MH literacy 3 nước châu Á — là tài liệu tham chiếu QUÝ cho VN khi xây dựng chương trình MH literacy vào curriculum. Philippines + Indonesia đều là LMIC giống VN → có thể adapt nhiều. Nhật Bản là HIC → reference cho mô hình lý tưởng. Cần đọc chi tiết Results + Discussion để rút ra khuyến nghị VN cụ thể.',
        'future': 'Nghiên cứu tương tự cho Việt Nam: phân tích chương trình "Giáo dục sức khỏe" hiện tại trong SGK phổ thông VN. So sánh với 3 nước Shibuya. Đề xuất lồng ghép MH literacy vào môn Giáo dục công dân + Sinh học.',
        'rating': 4,
        'notes_for_rag': 'So sánh chính sách school MH 3 nước Á (Philippines, Indonesia, Nhật) 2000-2023. Tài liệu tham chiếu cho VN curriculum.',
    },
    {
        'id': 'QT020',
        'descriptor': 'Liu_CBT_Delivery_GAD_NMA_2025',
        'pdf_src': 'CBT_Delivery_GAD_TranslPsych_2025.pdf',
        'pdf_dst': 'QT020_Liu_CBT_Delivery_GAD_NMA_2025.pdf',
        'folder': 'The-gioi_Khac',
        'title_vn': 'Các hình thức triển khai CBT cho rối loạn lo âu lan tỏa (GAD): Tổng quan hệ thống và phân tích tổng hợp mạng (NMA)',
        'title_en': 'CBT treatment delivery formats for generalized anxiety disorder: a systematic review and network meta-analysis of randomized controlled trials',
        'authors': 'Liu S, Xiao H, Duan Y, Shi L, Wang P, Cao L, Li H, Huang X, Qiu C',
        'affiliation': 'Trung Quốc (Sichuan University)',
        'journal': 'Translational Psychiatry (2025) · Nature Group',
        'year': 2025,
        'sample': '52 RCTs, N = 4.361 bệnh nhân GAD (tuổi TB 43, nữ 69,7 %)',
        'location': 'Toàn cầu (tìm đến tháng 9/2023)',
        'tool': 'Bayesian/frequentist NMA, random-effects, pairwise + network',
        'design': 'Systematic review + network meta-analysis',
        'outcomes': [
            'Individual CBT > Remote CBT: SMD 0,96 [95% CI: 0,13; 1,79]',
            'Individual CBT > Treatment-as-usual (TAU): SMD 1,12 [0,24; 2,00]',
            'Individual CBT > Waitlist: SMD 1,62 [1,03; 2,22]',
            'Group CBT (số liệu chi tiết ở chương Results — cần đọc full)',
            '75 % RCTs có low/unclear risk of bias',
        ],
        'strengths': [
            'NMA 52 RCT — cỡ bằng chứng LỚN cho GAD',
            'Transl Psychiatry — tạp chí Nature Q1',
            'So sánh nhiều format CBT: individual, group, remote, self-help',
            'Clear effect sizes với CI',
        ],
        'limitations': [
            '⚠ SCOPE KHÁC: đối tượng là NGƯỜI LỚN (tuổi TB 43) — KHÔNG PHẢI VTN. Áp dụng cho HS VN cần ngoại suy thận trọng',
            'Nữ 69,7 % — không cân bằng giới',
            'Remote CBT có SMD rộng 0,13-1,79 → heterogeneity cao',
            'Tìm kiếm đến 9/2023 — một số RCT mới hơn có thể chưa được gộp',
        ],
        'critique_short': 'NMA lớn về CBT FORMAT — nhưng đối tượng là NGƯỜI LỚN GAD (TB 43 tuổi), KHÔNG PHẢI VTN. Kết quả "Individual > Remote > TAU > Waitlist" là hữu ích làm reference hierarchy, nhưng effect size (SMD 0,96-1,62) KHÔNG THỂ áp dụng trực tiếp cho HS/VTN vì: (1) tâm lý phát triển khác; (2) động lực trị liệu khác (VTN ít tự giác hơn); (3) Individual CBT với VTN cần involve phụ huynh. Cần NMA tương tự nhưng restrict to < 18 tuổi.',
        'future': 'Lặp lại NMA nhưng chỉ với RCT ở trẻ em/VTN (< 18). Có thể so sánh với QT029 Li 2025 (NMA 30 RCT trẻ em/VTN). Phân tích subgroup format theo tuổi.',
        'rating': 3,
        'notes_for_rag': 'Người lớn GAD. Individual CBT > Remote > TAU > Waitlist (SMD 0,96-1,62). KHÔNG phải VTN — dùng tham chiếu hierarchy, không copy effect size.',
    },
]

# Build helpers (reuse from VN step)
def shade_cell(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)

def build_summary_docx(paper, out_path):
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    t = d.add_heading(f"{paper['id']} — Tóm tắt bài nghiên cứu", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)

    meta_tbl = d.add_table(rows=7, cols=2)
    meta_tbl.style = 'Table Grid'
    meta_rows = [
        ('Tên công trình (VN)', paper['title_vn']),
        ('Tên công trình (EN)', paper['title_en']),
        ('Tác giả', paper['authors']),
        ('Đơn vị', paper['affiliation']),
        ('Nơi công bố', paper['journal']),
        ('Năm', str(paper['year'])),
        ('Địa bàn + mẫu + công cụ', f"{paper['location']} | {paper['sample']} | Công cụ: {paper['tool']}"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        cell0 = meta_tbl.rows[i].cells[0]
        shade_cell(cell0, 'D9E1F2')
        p0 = cell0.paragraphs[0]
        r0 = p0.add_run(k)
        r0.bold = True
        meta_tbl.rows[i].cells[1].text = v
    d.add_paragraph()

    for sec, items, color in [
        ('Phương pháp nghiên cứu', [paper['design'], f"Mẫu: {paper['sample']}", f"Công cụ: {paper['tool']}"], (31, 73, 125)),
        ('Kết quả nghiên cứu', paper['outcomes'], (31, 73, 125)),
        ('Điểm mạnh', paper['strengths'], (54, 95, 44)),
        ('Hạn chế', paper['limitations'], (192, 80, 77)),
    ]:
        h = d.add_heading(sec, level=1)
        for r in h.runs:
            r.font.color.rgb = RGBColor(*color)
        for it in items:
            d.add_paragraph('• ' + it)

    h = d.add_heading('Phản biện ngắn', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(192, 0, 0)
    p = d.add_paragraph()
    r = p.add_run(paper['critique_short'])
    r.font.color.rgb = RGBColor(192, 0, 0)

    h = d.add_heading('Hướng nghiên cứu tiếp theo', level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(31, 73, 125)
    d.add_paragraph(paper['future'])

    stars = '⭐' * paper['rating']
    h = d.add_heading(f"Đánh giá chất lượng: {stars} ({paper['rating']}/5)", level=1)
    for r in h.runs:
        r.font.color.rgb = RGBColor(191, 97, 14)

    if paper.get('notes_for_rag'):
        d.add_paragraph()
        p = d.add_paragraph()
        r = p.add_run('Key fact cho RAG: ')
        r.bold = True
        r.font.size = Pt(10)
        p.add_run(paper['notes_for_rag']).font.size = Pt(10)

    d.save(out_path)

# Execute
log = []
with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

for paper in PAPERS:
    src = QT_FOLDER_KHAC / paper['pdf_src']
    dst = QT_FOLDER_KHAC / paper['pdf_dst']
    if not src.exists():
        log.append(f"SKIP {paper['id']}: PDF source {paper['pdf_src']} missing")
        continue
    if not dst.exists():
        src.rename(dst)
        log.append(f"RENAMED {paper['id']}: {src.name} → {dst.name}")
    else:
        log.append(f"SKIP rename {paper['id']}: dst exists")

    tt_name = f"{paper['id']}_{paper['descriptor']}.docx"
    tt_path = TT_DIR / tt_name
    build_summary_docx(paper, tt_path)
    log.append(f"SUMMARY {paper['id']}: {tt_name} ({tt_path.stat().st_size//1024} KB)")

    idx[paper['id']] = {
        'id': paper['id'],
        'descriptor': paper['descriptor'],
        'summary': tt_name,
        'translation': None,
        'pdf': paper['pdf_dst'],
        'pdf_folder': paper['folder'],
    }
    log.append(f"INDEX {paper['id']}: added")

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

n_vn = sum(1 for k in idx if k.startswith('VN'))
n_qt = sum(1 for k in idx if k.startswith('QT'))
log.append(f'\nTotal entries: {len(idx)} ({n_vn} VN + {n_qt} QT)')

for l in log:
    print(l)
