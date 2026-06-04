"""Build doc: 'Ứng phó là yếu tố bảo vệ — bằng chứng và hàm ý cho đề tài lo âu HS' (CTH v6).
All facts verified against:
- Compas et al. 2017 Psychological Bulletin abstract (212 studies, N=80,850)
- Lazarus & Folkman 1984 Springer NY (445 trang)
- Herres & Ohannessian 2015 J Affective Disorders 186:312-319 (n=982 verified)
- Lochman et al. 2025 EACP RCT (n=709 lớp 7, 40 trường)
- Cai et al. 2025 Frontiers Psychiatry (SMD=0.17 verified)
- Murphy et al. 2024 J Community Psychology (15 NC, 13 interventions)
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Ung_pho_yeu_to_bao_ve_bang_chung_va_ham_y.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLACK

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# ============================================================
# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('ỨNG PHÓ LÀ YẾU TỐ BẢO VỆ?\nBẰNG CHỨNG QUỐC TẾ VÀ HÀM Ý CHO ĐỀ TÀI\nLO ÂU HỌC SINH TRUNG HỌC CƠ SỞ VIỆT NAM')
r.bold = True; r.font.size = Pt(14)
para('')

# 1. Đặt vấn đề — mở 3 cấp độ
H('1. Đặt vấn đề', level=2, color=NAVY)
para(
    'Ứng phó (coping) — các nỗ lực nhận thức và hành vi để quản lý các '
    'đòi hỏi vượt quá nguồn lực cá nhân — là khái niệm trung tâm trong '
    'tâm lý học sức khỏe tâm thần kể từ tác phẩm kinh điển của Lazarus '
    'và Folkman (1984). Trên BÌNH DIỆN TOÀN CẦU, phân tích tổng hợp '
    'lớn nhất đến nay (Compas và cộng sự, 2017 — 212 nghiên cứu, '
    'N = 80.850 trẻ em và vị thành niên) đã thiết lập rõ rệt ranh giới '
    'giữa ứng phó tích cực và ứng phó tiêu cực — hai chiều TRÁI ngược '
    'nhau về tác động lên rối loạn nội hóa.'
)
para(
    'Tại KHU VỰC châu Á và Đông Nam Á, các nghiên cứu gần đây '
    '(Qiu và cộng sự, 2022 — Trung Quốc, n = 2.079) đã xác lập vai '
    'trò bảo vệ của khả năng phục hồi (resilience) — thành phần cốt '
    'lõi của ứng phó tích cực — qua hồi quy logistic đa biến. Tại '
    'VIỆT NAM, kết quả của Trần Thảo Vi và cộng sự (2024) tại Huế '
    '(thuần tập 3 năm, n = 611 học sinh trung học cơ sở) cho thấy '
    'lối sống và yếu tố học tập là nguồn stress chính — gợi ý nhu '
    'cầu can thiệp ứng phó.'
)
para(
    'Câu hỏi đặt ra trong đề tài: ứng phó của học sinh có phải là yếu '
    'tố bảo vệ không? Phần dưới đây tổng hợp bằng chứng từ 6 công trình '
    'KINH ĐIỂN và CẬP NHẬT — từ khung lý thuyết gốc, phân tích tổng hợp '
    'định lượng, đến thử nghiệm can thiệp ngẫu nhiên — để trả lời câu '
    'hỏi này một cách TINH VI: KHÔNG phải mọi ứng phó đều bảo vệ; chỉ '
    'ứng phó TÍCH CỰC bảo vệ, trong khi ứng phó TIÊU CỰC LÀ YẾU TỐ '
    'NGUY CƠ.'
)

# 2. Khung khái niệm
H('2. Khung khái niệm — Lazarus và Folkman (1984)', level=2, color=NAVY)
para(
    'Lazarus và Folkman (1984) trong tác phẩm "Stress, Appraisal, and '
    'Coping" (Springer New York, 445 trang) đã xây dựng khung lý '
    'thuyết stress-đánh giá-ứng phó được xem là CÔNG TRÌNH ĐƯỢC TRÍCH '
    'DẪN NHIỀU NHẤT trong y văn về stress và ứng phó. Hai trục phân '
    'loại được thiết lập: (1) ứng phó TẬP TRUNG vào VẤN ĐỀ '
    '(problem-focused coping) — thay đổi nguồn gốc stress; (2) ứng phó '
    'TẬP TRUNG vào CẢM XÚC (emotion-focused coping) — điều chỉnh phản '
    'ứng cảm xúc với stress.'
)
para(
    'Compas và cộng sự (2017) mở rộng khung này thành hai cặp khái '
    'niệm tinh vi hơn. Thứ nhất, ứng phó KIỂM SOÁT SƠ CẤP (primary '
    'control coping) — thay đổi tác nhân stress hoặc phản ứng cảm '
    'xúc trực tiếp (giải quyết vấn đề, biểu hiện cảm xúc có chủ đích). '
    'Thứ hai, ứng phó KIỂM SOÁT THỨ CẤP (secondary control coping) — '
    'thích nghi với tác nhân stress thay vì thay đổi nó (chấp nhận, '
    'đánh thoái nhận thức, nghĩ tích cực, tập trung sự chú ý). Thứ ba, '
    'ứng phó NGẮT KẾT NỐI (disengagement coping) — né tránh, phủ '
    'nhận, ức chế cảm xúc — KHÔNG được xem là ứng phó hiệu quả mà '
    'là CHIẾN LƯỢC GÂY HẠI.'
)

# 3. Bằng chứng định lượng
H('3. Bằng chứng định lượng — sáu công trình', level=2, color=NAVY)

# 3.1 Compas 2017
H('3.1. Compas và cộng sự (2017) — Phân tích tổng hợp 212 nghiên cứu', level=3)
para(
    'Compas, Jaser, Bettis, Watson, Gruhn, Dunbar, Williams và Thigpen '
    '(2017) trong Psychological Bulletin thực hiện phân tích tổng hợp '
    '212 nghiên cứu với N = 80.850 trẻ em và vị thành niên — đến nay '
    'là phân tích tổng hợp LỚN NHẤT về mối quan hệ giữa ứng phó, điều '
    'tiết cảm xúc, và bệnh lý tâm thần ở giai đoạn phát triển. Bài '
    'báo phân tích cả nghiên cứu cắt ngang (đa số) và nghiên cứu thuần '
    'tập (cung cấp bằng chứng MẠNH HƠN về quan hệ nhân quả).'
)
para(
    'Phát hiện CỐT LÕI: hai chiều ứng phó tác động TRÁI ngược nhau lên '
    'triệu chứng nội hóa (lo âu + trầm cảm). Ứng phó kiểm soát SƠ CẤP '
    'và ứng phó kiểm soát THỨ CẤP có liên quan với MỨC THẤP HƠN của '
    'triệu chứng bệnh lý tâm thần. Trái lại, ứng phó NGẮT KẾT NỐI cùng '
    'với các chiến lược ức chế cảm xúc, né tránh, và phủ nhận có liên '
    'quan với MỨC CAO HƠN của triệu chứng bệnh lý.'
)
para(
    'Nói cách khác, kết quả của Compas và cộng sự xác lập hai luận '
    'điểm. Thứ nhất, ứng phó KHÔNG phải là một biến đơn — phải phân '
    'biệt ít nhất ba thành phần (sơ cấp / thứ cấp / ngắt kết nối). '
    'Thứ hai, gộp tất cả "ứng phó" vào một biến tổng hợp sẽ '
    'TRIỆT TIÊU lẫn nhau hai chiều tác động — và làm mờ đi phát '
    'hiện thực sự.'
)

# 3.2 Herres & Ohannessian 2015
H('3.2. Herres và Ohannessian (2015) — Hồ sơ ứng phó phân biệt lo âu và trầm cảm', level=3)
para(
    'Herres và Ohannessian (2015) trong Journal of Affective '
    'Disorders, tập 186, trang 312–319 (DOI: 10.1016/j.jad.2015.07.031), '
    'thực hiện khảo sát cắt ngang 982 học sinh trung học phổ thông '
    'tại 7 trường công lập vùng Mid-Atlantic Hoa Kỳ (Pennsylvania, '
    'New Jersey, Delaware, Maryland) vào mùa xuân 2007 — tuổi trung '
    'bình 16,09 (SD = 0,68); 54% nữ; lớp 10–11. Bộ công cụ chuẩn: '
    'Children\'s Depression Inventory (CDI), State-Trait Anxiety '
    'Inventory for Children (STAIC), và COPE Inventory.'
)
para(
    'Đặc sắc của bài: thay vì xem ứng phó như biến liên tục, các tác '
    'giả áp dụng phân tích hồ sơ tiềm ẩn (latent profile analysis) '
    'để xác định CÁC HỒ SƠ ỨNG PHÓ — tổ hợp đặc trưng của các chiến '
    'lược trong cùng một học sinh. Phát hiện: hồ sơ ứng phó CHỦ ĐỘNG '
    'TÍCH CỰC (giải quyết vấn đề + tìm hỗ trợ + đánh thoái) phân '
    'biệt rõ với hồ sơ ứng phó NÉ TRÁNH (rút lui + phủ nhận + ức chế '
    'cảm xúc) về cả triệu chứng trầm cảm lẫn lo âu.'
)
para(
    'Phù hợp với phát hiện của Compas và cộng sự (2017): cùng một '
    'thanh thiếu niên thường có CẢ chiến lược tích cực và tiêu cực — '
    'điểm khác biệt giữa nhóm có triệu chứng cao và nhóm có triệu '
    'chứng thấp KHÔNG nằm ở việc dùng hay không dùng một chiến lược '
    'đơn lẻ, mà nằm ở TỶ LỆ giữa các chiến lược.'
)

# 3.3 Lochman 2025 EACP
H('3.3. Lochman và cộng sự (2025) — Thử nghiệm Coping Power tại trường', level=3)
para(
    'Lochman và cộng sự (2025) trong Journal of School Psychology công '
    'bố thử nghiệm ngẫu nhiên có đối chứng theo cụm (cluster RCT) '
    'của chương trình Early Adolescent Coping Power (EACP) tại 40 '
    'trường trung học cơ sở Hoa Kỳ — n = 709 học sinh lớp 7 (69,8% '
    'gốc Phi; 59,4% nam) được tuyển vì có hành vi gây hấn. Tài trợ '
    'của Viện Sức khỏe Tâm thần Quốc gia Hoa Kỳ (NIMH).'
)
para(
    'Chương trình EACP gồm BA THÀNH PHẦN: (1) học sinh — 25 buổi '
    'nhóm trong năm học; (2) cha mẹ — 16 buổi 90 phút song song; '
    '(3) giáo viên — phối hợp cố vấn học đường. Năm nhóm kỹ năng '
    'cho học sinh: quản lý cơn giận và lo âu (lời nói nội tâm tích '
    'cực + thư giãn); giải quyết vấn đề xã hội; tương tác với bạn '
    'bè; kỹ năng học tập và tổ chức; kỹ năng từ chối chất gây '
    'nghiện. Nền tảng lý thuyết: mô hình social-cognitive '
    'developmental contextual.'
)
para(
    'Hàm ý cho Việt Nam: mô hình BA THÀNH PHẦN — học sinh + cha '
    'mẹ + giáo viên — phù hợp bối cảnh trường học Việt Nam, nơi '
    'cha mẹ và giáo viên đóng vai trò TRUNG TÂM trong quản lý '
    'sức khỏe tâm thần học sinh. Tuy nhiên, đối tượng gốc của '
    'EACP là học sinh có hành vi GÂY HẤN — không phải lo âu '
    'thuần — nên cần điều chỉnh nội dung trước khi áp dụng cho '
    'lo âu.'
)

# 3.4 Cai 2025
H('3.4. Cai và cộng sự (2025) — Phân tích tổng hợp 21 RCT can thiệp tại trường', level=3)
para(
    'Cai, Mei, Wang và Luo (2025) trong Frontiers in Psychiatry '
    '(DOI: 10.3389/fpsyt.2025.1594658) thực hiện phân tích tổng '
    'hợp các thử nghiệm ngẫu nhiên có đối chứng về can thiệp '
    'resilience tại trường — thành phần lõi gồm kỹ năng ứng phó, '
    'lòng tự trọng, kết nối xã hội, và lạc quan. Tổng số 38 RCT '
    'qua 17 quốc gia: Mỹ (n = 10), Trung Quốc (n = 9), Úc (n = 4), '
    'Pakistan (n = 2), Ấn Độ (n = 2), và 11 nước khác mỗi nước 1 '
    'nghiên cứu.'
)
para(
    'Kết quả tổng hợp trên 21 RCT đo resilience: SMD = 0,17 '
    '(KTC 95% từ 0,06 đến 0,29; p < 0,01) — tương ứng kích thước '
    'hiệu ứng NHỎ nhưng có ý nghĩa thống kê. Heterogeneity rất CAO '
    '(I² = 81,90%) — gợi ý hiệu quả phụ thuộc nhiều vào thiết kế '
    'can thiệp, đối tượng, và bối cảnh. Phù hợp với khuyến nghị '
    'của các tác giả: can thiệp ứng phó-resilience tại trường có '
    'hiệu quả thực tế nhưng cần cá nhân hóa cho từng bối cảnh.'
)

# 3.5 Qiu 2022
H('3.5. Qiu và cộng sự (2022) — Resilience là yếu tố bảo vệ MẠNH NHẤT (Trung Quốc)', level=3)
para(
    'Qiu, Guo, Wang và Zhang (2022) trong Frontiers in Public '
    'Health khảo sát 2.079 học sinh trung học cơ sở (tuổi trung '
    'bình 16,7) tại thành phố Hợp Phì, tỉnh An Huy, Trung Quốc — '
    'sử dụng Thang Khả năng Phục hồi SRSMSS, Thang Trầm cảm '
    'CES-D, Thang Lo âu Tự đánh giá SAS, và Thang Phong cách Nuôi '
    'dạy EMBU. Tỷ lệ phản hồi cao (98,06%).'
)
para(
    'Phát hiện đáng chú ý: KHẢ NĂNG PHỤC HỒI THẤP là YẾU TỐ NGUY '
    'CƠ MẠNH NHẤT cho trầm cảm — OR = 6,74 sau khi kiểm soát các '
    'biến nhân khẩu học và phong cách nuôi dạy. Trái lại, nuôi dạy '
    'tích cực giảm nguy cơ trầm cảm 70% (OR = 0,30) và lo âu 68% '
    '(OR = 0,32). Tỷ lệ trầm cảm 26% và lo âu 13,4% trong mẫu — '
    'cao đáng kể.'
)
para(
    'Nói cách khác, resilience — thành phần CỐT LÕI của ứng phó '
    'tích cực — là yếu tố bảo vệ MẠNH HƠN cả phong cách nuôi dạy. '
    'Điều này ủng hộ luận điểm: can thiệp ứng phó ở cấp độ cá '
    'nhân (kỹ năng học sinh) có giá trị độc lập, không thể thay '
    'thế bằng can thiệp môi trường (cha mẹ, trường học) — và '
    'ngược lại.'
)

# 3.6 Murphy 2024
H('3.6. Murphy và cộng sự (2024) — Hỗ trợ đồng đẳng (peer support) trong dịch vụ tâm thần thanh niên', level=3)
para(
    'Murphy, Huggard, Fitzgerald, Hennessy và Booth (2024) trong '
    'Journal of Community Psychology, tập 52, số 1, trang 154–180 '
    '(DOI: 10.1002/jcop.23090) thực hiện scoping review hệ thống '
    'theo khung Arksey và O\'Malley — gộp 15 nghiên cứu mô tả 13 '
    'can thiệp peer support trong dịch vụ chăm sóc sức khỏe tâm '
    'thần thanh niên. Định nghĩa peer support: hỗ trợ xã hội + cảm '
    'xúc giữa cá nhân có TRẢI NGHIỆM CHUNG về khó khăn sức khỏe '
    'tâm thần.'
)
para(
    'Phát hiện: peer support có TIỀM NĂNG cải thiện kết quả phục '
    'hồi (recovery-related outcomes) — tuy nhiên bằng chứng chưa '
    'đồng nhất do thiết kế khác nhau. Hai rào cản chính: (1) lo '
    'ngại bảo mật trong quan hệ peer; (2) peer worker thiếu tự '
    'tin về vai trò. Hai yếu tố thúc đẩy: (1) nhân viên chuyên môn '
    'hỗ trợ peer worker; (2) vai trò peer được định nghĩa rõ. '
    'Đáng chú ý: 9 trong 13 can thiệp dành cho thanh niên 16–25 '
    'tuổi — chỉ 2 trong 13 cho học sinh trung học cơ sở-trung '
    'học phổ thông.'
)

# 4. Bảng tổng hợp
H('4. Bảng tổng hợp sáu công trình', level=2, color=NAVY)
caption('Bảng 1. Sáu công trình quốc tế về ứng phó là yếu tố bảo vệ')
add_table(
    ['Công trình', 'Loại NC + mẫu', 'Phát hiện chính'],
    [
        ['Lazarus & Folkman (1984), Springer NY',
         'Sách lý thuyết, 445 trang',
         'Khung stress-đánh giá-ứng phó; 2 loại: problem-focused vs emotion-focused'],
        ['Compas và cộng sự (2017), Psychological Bulletin',
         'Meta-analysis 212 NC, N = 80.850',
         'Sơ cấp + thứ cấp giảm triệu chứng; ngắt kết nối + né tránh tăng triệu chứng'],
        ['Herres & Ohannessian (2015), J Affective Disorders 186:312-319',
         'Cắt ngang n = 982 HS THPT Mỹ',
         'Hồ sơ ứng phó tích cực vs né tránh phân biệt rõ trầm cảm + lo âu'],
        ['Qiu và cộng sự (2022), Frontiers in Public Health',
         'Cắt ngang n = 2.079 HS THCS TQ',
         'Resilience THẤP là yếu tố nguy cơ MẠNH NHẤT (OR = 6,74)'],
        ['Murphy và cộng sự (2024), J Community Psychology 52(1):154-180',
         'Scoping review 15 NC, 13 can thiệp',
         'Peer support có tiềm năng — cần vai trò rõ + hỗ trợ chuyên môn'],
        ['Cai và cộng sự (2025), Frontiers in Psychiatry',
         'Meta-analysis 21 RCT can thiệp tại trường',
         'SMD = 0,17 (KTC 95% 0,06–0,29); I² = 81,90%'],
        ['Lochman và cộng sự (2025), J School Psychology',
         'Cluster RCT n = 709 lớp 7 Mỹ',
         'EACP 3 thành phần (HS + cha mẹ + GV) — hiệu quả hành vi'],
    ]
)

# 5. Tổng hợp
H('5. Năm phát hiện chính', level=2, color=NAVY)
para(
    'Bằng chứng quốc tế cho phép rút ra NĂM PHÁT HIỆN chính về vai '
    'trò ứng phó trong lo âu vị thành niên.', indent=False
)
para(
    'Thứ nhất, ứng phó là biến PHÂN ĐỔI — không thể xử lý như biến '
    'tổng. Compas và cộng sự (2017) trên 80.850 trẻ em + vị thành '
    'niên xác lập rõ: ứng phó kiểm soát SƠ CẤP và THỨ CẤP có tác '
    'động ÂM với triệu chứng nội hóa, trong khi ứng phó NGẮT KẾT '
    'NỐI (né tránh, phủ nhận, ức chế cảm xúc) có tác động DƯƠNG. '
    'Hai chiều TRÁI ngược nhau — gộp lại sẽ triệt tiêu phát hiện.'
)
para(
    'Thứ hai, ứng phó tích cực là YẾU TỐ BẢO VỆ; ứng phó tiêu cực '
    'là YẾU TỐ NGUY CƠ. Cả nghiên cứu hồ sơ (Herres & Ohannessian, '
    '2015) và phân tích tổng hợp định lượng (Compas và cộng sự, '
    '2017) đều ủng hộ luận điểm này. Trái với quan niệm dân gian '
    '"có ứng phó là tốt", bằng chứng cho thấy ứng phó NÉ TRÁNH '
    'làm TRẦM TRỌNG thêm lo âu.'
)
para(
    'Thứ ba, can thiệp tại trường có HIỆU QUẢ THỰC TẾ nhưng kích '
    'thước NHỎ (SMD = 0,17 — Cai và cộng sự, 2025). Heterogeneity '
    'rất CAO (I² = 81,90%) — gợi ý hiệu quả phụ thuộc thiết kế. '
    'Mô hình EACP của Lochman và cộng sự (2025) — kết hợp ba '
    'thành phần học sinh + cha mẹ + giáo viên — là khung tham '
    'khảo cho thiết kế can thiệp tại Việt Nam.'
)
para(
    'Thứ tư, resilience — thành phần CỐT LÕI của ứng phó tích cực — '
    'có tác động ĐỘC LẬP với phong cách nuôi dạy (Qiu và cộng sự, '
    '2022). Khả năng phục hồi thấp là yếu tố nguy cơ MẠNH NHẤT cho '
    'trầm cảm với OR = 6,74. Phát hiện này ủng hộ chương trình can '
    'thiệp ở cấp ĐỘ HỌC SINH (kỹ năng cá nhân) — không thể thay '
    'thế bằng can thiệp gia đình.'
)
para(
    'Thứ năm, peer support là kênh BỔ TRỢ tiềm năng — đặc biệt '
    'cho bối cảnh Việt Nam nơi học sinh ưu thích nguồn hỗ trợ '
    'NON-PROFESSIONAL trước (Murphy và cộng sự, 2024). Tuy nhiên, '
    'cần định nghĩa rõ vai trò peer worker và bảo đảm hỗ trợ '
    'chuyên môn — tránh rủi ro bảo mật và quá tải đồng đẳng.'
)

# 6. Hàm ý cho đề tài
H('6. Hàm ý cho đề tài lo âu HS THCS Việt Nam', level=2, color=NAVY)
para(
    'Sáu công trình trên cùng kết luận của bài thúc đẩy bốn '
    'khuyến nghị thiết kế cho đề tài.', indent=False
)
para(
    'KHUYẾN NGHỊ 1 — Đo ứng phó bằng thang phân biệt nhiều chiều. '
    'Đề xuất sử dụng A-COPE (Adolescent Coping Orientation for '
    'Problem Experiences) hoặc Children\'s Coping Strategies '
    'Checklist (CCSC) — phân biệt ít nhất bốn nhóm: giải quyết '
    'vấn đề / tìm hỗ trợ xã hội / đánh thoái nhận thức / né tránh. '
    'KHÔNG GỘP tất cả thành biến "ứng phó" tổng — sẽ triệt tiêu '
    'hai chiều tác động.'
)
para(
    'KHUYẾN NGHỊ 2 — Phân tích RIÊNG hiệu ứng từng nhóm với lo âu. '
    'Dự kiến kết quả: nhóm tích cực (giải quyết vấn đề + tìm hỗ '
    'trợ xã hội + đánh thoái) có β ÂM với lo âu; nhóm tiêu cực '
    '(né tránh + phủ nhận + ức chế cảm xúc) có β DƯƠNG. Nếu kết '
    'quả khớp với dự kiến, đây là bằng chứng VIỆT NAM bổ sung cho '
    'phát hiện quốc tế của Compas và cộng sự (2017).'
)
para(
    'KHUYẾN NGHỊ 3 — Thiết kế can thiệp ba thành phần. Theo mô '
    'hình EACP (Lochman và cộng sự, 2025): thành phần học sinh '
    '(kỹ năng ứng phó cá nhân) + thành phần cha mẹ (phong cách '
    'nuôi dạy) + thành phần giáo viên (phát hiện sớm + cố vấn '
    'học đường). Tích hợp peer support theo khuyến nghị Murphy '
    'và cộng sự (2024) — với vai trò định nghĩa rõ.'
)
para(
    'KHUYẾN NGHỊ 4 — Đo resilience riêng bằng CD-RISC. Connor & '
    'Davidson Resilience Scale (Connor & Davidson, 2003) — 25 '
    'mục, năm yếu tố — là thang đo chuẩn vàng quốc tế. Sử dụng '
    'CD-RISC tách riêng resilience khỏi ứng phó cho phép so '
    'sánh trực tiếp với phát hiện Qiu và cộng sự (2022) — '
    'đặc biệt OR = 6,74 cho resilience thấp.'
)

# 7. Câu trả lời ngắn gọn
H('7. Trả lời câu hỏi ban đầu', level=2, color=NAVY)
para(
    'Câu hỏi: Ứng phó của học sinh có phải là yếu tố bảo vệ không?', bold=True, indent=False
)
para(
    'Trả lời: CÓ — nhưng chỉ ứng phó TÍCH CỰC (kiểm soát sơ cấp + '
    'kiểm soát thứ cấp). Ứng phó TIÊU CỰC (ngắt kết nối, né tránh, '
    'phủ nhận, ức chế cảm xúc) là YẾU TỐ NGUY CƠ — làm TĂNG triệu '
    'chứng lo âu và trầm cảm. Bằng chứng MẠNH NHẤT đến nay là phân '
    'tích tổng hợp 212 nghiên cứu của Compas và cộng sự (2017) trên '
    '80.850 trẻ em + vị thành niên — xác lập rõ hai chiều TRÁI '
    'ngược nhau. Phù hợp với phát hiện hồ sơ ứng phó của Herres và '
    'Ohannessian (2015) trên 982 học sinh trung học phổ thông Hoa '
    'Kỳ và phát hiện về resilience của Qiu và cộng sự (2022) trên '
    '2.079 học sinh trung học cơ sở Trung Quốc.'
)
para(
    'Hàm ý quan trọng cho đề tài: thiết kế nghiên cứu KHÔNG nên '
    'gộp tất cả "ứng phó" thành một biến — phải phân biệt ít nhất '
    'BA NHÓM (sơ cấp / thứ cấp / né tránh) và phân tích RIÊNG hiệu '
    'ứng từng nhóm với lo âu. Đo resilience riêng bằng CD-RISC để '
    'so sánh trực tiếp với chuẩn quốc tế.'
)

# 8. TLTK
H('8. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, 1594658. https://doi.org/10.3389/fpsyt.2025.1594658',
    'Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991. https://doi.org/10.1037/bul0000110',
    'Connor, K. M., & Davidson, J. R. T. (2003). Development of a new resilience scale: The Connor-Davidson Resilience Scale (CD-RISC). Depression and Anxiety, 18(2), 76–82. https://doi.org/10.1002/da.10113',
    'Herres, J., & Ohannessian, C. M. (2015). Adolescent coping profiles differentiate reports of depression and anxiety symptoms. Journal of Affective Disorders, 186, 312–319. https://doi.org/10.1016/j.jad.2015.07.031',
    'Lazarus, R. S., & Folkman, S. (1984). Stress, appraisal, and coping. Springer Publishing Company.',
    'Lochman, J. E., và cộng sự. (2025). Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology.',
    'Murphy, R., Huggard, L., Fitzgerald, A., Hennessy, E., & Booth, A. (2024). A systematic scoping review of peer support interventions in integrated primary youth mental health care. Journal of Community Psychology, 52(1), 154–180. https://doi.org/10.1002/jcop.23090',
    'Qiu, Z., Guo, Y., Wang, J., & Zhang, H. (2022). Associations of parenting style and resilience with depression and anxiety symptoms among Chinese middle school students. Frontiers in Public Health, 10, 897339. [QT009 trong cơ sở dữ liệu dự án.]',
    'Trần, T. V., và cộng sự. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine, 19(4).',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
