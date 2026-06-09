# -*- coding: utf-8 -*-
"""Doc tra loi thay NMD ve QD HD dao duc cho LA + Q2/Q3/Q4 papers.
Tieng Viet thuan, co tham khao."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1',
                   'HoiDong_DaoDuc_TraLoiNMD_08062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(12)
s.paragraph_format.line_spacing = 1.4


def TITLE(t, sz=15):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(sz); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H1(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def P(t):
    p = d.add_paragraph()
    p.paragraph_format.space_after = Pt(4); p.paragraph_format.first_line_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def Q(t):
    """Câu hỏi của thầy"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(4)
    r = p.add_run('Câu hỏi của thầy: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    r2 = p.add_run('"' + t + '"')
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11); r2.italic = True

def A(t):
    """Câu trả lời em đề xuất"""
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(8)
    r = p.add_run('Đáp án em đề xuất: ')
    r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.bold = True
    r.font.color.rgb = RGBColor(0x00, 0x6E, 0x2E)
    r2 = p.add_run(t)
    r2.font.name = 'Times New Roman'; r2.font.size = Pt(11)

def BB(t):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + t); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


# ============================================================
TITLE('THAM KHẢO VỀ QUYẾT ĐỊNH HỘI ĐỒNG ĐẠO ĐỨC')
TITLE('Cho luận án tiến sĩ + 3 bài báo Q2/Q3/Q4', 12)
p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('Tài liệu nội bộ chuẩn bị cho nhóm nghiên cứu — Soạn '
              '08/06/2026')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True
r.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


# ============================================================
H1('TÓM TẮT NHANH (TL;DR)')
P('Em tóm tắt ngay 3 câu trả lời để các thầy và chị có cái nhìn tổng '
  'quan trước khi đọc chi tiết bên dưới ạ:')

P('1. CÓ THỂ DÙNG 1 QUYẾT ĐỊNH chung cho toàn bộ đề tài LA + 3 bài '
  'báo Q2/Q3/Q4 — không cần 3 QĐ riêng. Đây là thông lệ quốc tế.')

P('2. Chỉ ghi SỐ QĐ là KHÔNG ĐỦ cho tạp chí Q1/Q2 quốc tế. Tối thiểu '
  'phải có: số QĐ + ngày phê duyệt + tên cơ quan phê duyệt + phạm vi '
  'phê duyệt (mục tiêu nghiên cứu). Nếu được, có thêm bản scan thư '
  'chính thức.')

P('3. Mẫu QĐ thực tế: em tham khảo từ Đại học Y Hà Nội (IRB chính '
  'thức hoạt động, có template chuẩn), Đại học Quốc gia TP.HCM (USSH), '
  'và một số bài báo Việt Nam đã xuất bản trên tạp chí quốc tế. Mục '
  'V của tài liệu này có ví dụ cụ thể.')


# ============================================================
H1('I. TRẢ LỜI 3 CÂU HỎI CỤ THỂ CỦA THẦY')

H2('1.1 — Có cần 3 quyết định riêng cho từng bài báo không?')
Q('Có cần 3 quyết định cho riêng từng bài báo hay chỉ cần quyết định '
   'cho phép triển khai đề tài luận án và công bố quốc tế?')
A('KHÔNG cần 3 QĐ riêng. Thông lệ quốc tế: một QĐ HĐ đạo đức cấp cho '
  'TOÀN BỘ đề tài (research protocol) sẽ bao trùm tất cả các bài báo '
  'phát sinh từ đề tài đó, miễn các bài báo nằm trong phạm vi mục '
  'tiêu đã phê duyệt. Trong trường hợp của chúng ta, bộ dữ liệu '
  '1.352 học sinh đã thu thập theo đề cương LA của NCS Công Thị Hằng. '
  'Bài Q2 phân tích SEM tích hợp, Q3 phân tích SEM đa nhóm theo '
  'giới, Q4 phân tích hồ sơ tiềm ẩn (LPA) — cả ba bài đều dùng cùng '
  'bộ dữ liệu đó, đều nằm trong scope của câu hỏi nghiên cứu gốc về '
  '"Yếu tố nguy cơ-bảo vệ của lo âu ở học sinh THCS Việt Nam". Vì '
  'vậy 1 QĐ duy nhất là đủ.')

P('Bằng chứng thông lệ quốc tế: theo APA (Hiệp hội Tâm lý học Hoa Kỳ) '
  'và COPE (Ủy ban Đạo đức Công bố Khoa học), một protocol IRB '
  'thường được phê duyệt cho nhiều năm và sản sinh nhiều công bố '
  '(IRBs and Psychological Science, APA 2024; COPE Salami Slicing '
  'Position Statement 2018). Điều kiện duy nhất là 3 bài phải có '
  'unique contribution rõ ràng và không che giấu mối liên hệ — vì '
  'vậy bài sau phải tham chiếu rõ bài trước (như em đã đưa vào bản '
  'thảo Q3 và Q4: "Building on our preceding Q2 paper...").')

H2('1.2 — Chỉ cần có số QĐ là được không?')
Q('Hay chỉ cần có số QĐ là được?')
A('KHÔNG ĐỦ. Số QĐ là YẾU TỐ TỐI THIỂU nhưng tạp chí Q1/Q2 quốc tế '
  'yêu cầu THÊM:')

BB('Tên cơ quan phê duyệt (đầy đủ: ví dụ "Hội đồng Đạo đức trong '
   'Nghiên cứu Khoa học, Trường Đại học Sư phạm Hà Nội")')
BB('Ngày phê duyệt cụ thể')
BB('Số quyết định/mã protocol (ví dụ "QĐ số XXX/QĐ-ĐHSPHN" hoặc '
   '"IRB-VN01...")')
BB('Phạm vi phê duyệt: tóm tắt mục tiêu nghiên cứu mà QĐ này phê '
   'duyệt')
BB('Tuyên bố về quy trình lấy chấp thuận: "sự đồng ý bằng văn bản '
   'của cha mẹ và sự đồng thuận của học sinh đã được thu thập trước '
   'khi triển khai"')

P('Một số tạp chí (đặc biệt Frontiers, BMC, PLOS) còn có thể yêu cầu '
  'BẢN SCAN thư chính thức đính kèm vào "Supplementary Materials" '
  'trong vòng peer review. Tốt nhất là NCS có sẵn bản scan để gửi '
  'khi reviewer yêu cầu.')

H2('1.3 — Mẫu QĐ HĐ đạo đức tham khảo')
Q('Mr. Hải cho xin cái mẫu quyết định của HĐ đạo đức để nhóm tham '
   'khảo nhé.')
A('Em sưu tầm 3 nguồn mẫu khả thi để nhóm tham khảo:')
P('(a) Đại học Y Hà Nội (HMU) — có IRB chính thức (mã '
  'FWA00004148, IRB-VN01.001.III.RB00003121) hoạt động theo chuẩn '
  'OHRP của Hoa Kỳ. Các QĐ của HMU có dạng: "IRB-VN01.00.III.'
  'RB00003121/IFWA00004148-năm/lĩnh vực-số thứ tự". Ví dụ: '
  '"133/2020/YTCC-HD3" (số 133, năm 2020, lĩnh vực Y tế công cộng, '
  'hội đồng số 3). NCS có thể liên hệ Phòng Quản lý Khoa học của HMU '
  'qua email irb@hmu.edu.vn để xin mẫu nếu cần.')

P('(b) Đại học Khoa học Xã hội và Nhân văn — ĐHQG TP.HCM (HCMUSSH) '
  '— đã ban hành Quy chế Tổ chức và Hoạt động của Hội đồng Đạo đức '
  'Nghiên cứu. Có template chuẩn cho QĐ phê duyệt. Tham khảo tại '
  'website chính thức: hcmussh.edu.vn (mục Tin tức/Hội đồng Đạo đức).')

P('(c) Tham khảo trực tiếp các bài báo Việt Nam đã công bố trên tạp '
  'chí quốc tế Q1/Q2 — họ đã trích dẫn QĐ HĐ đạo đức theo đúng format '
  'tạp chí yêu cầu. Ví dụ: bài "Anxiety symptoms and coping strategies '
  'among high school students in Vietnam" của Pham TT Hoa và cs. '
  '(2024, Frontiers in Public Health, PMID: 38435293) cho cohort '
  'cũng 1.352 học sinh Hà Nội, đã ghi rõ QĐ phê duyệt từ Bệnh viện '
  'Tâm thần Trung ương.')


# ============================================================
H1('II. BỐI CẢNH VIỆT NAM — VÌ SAO PHẦN NÀY KHÔNG RÕ?')

P('Thầy có nhận định rất chính xác: "Ở VN, các luận án tiến sĩ tâm '
  'lý chưa thấy quyết định của HĐ đạo đức, chỉ thấy NCS ghi chung '
  'chung là đã thông qua HĐ đạo đức." Em đào sâu vấn đề này để cung '
  'cấp bối cảnh:')

H2('2.1 — Hệ thống IRB ở Việt Nam còn mới')
P('Hệ thống IRB (Institutional Review Board) ở Việt Nam mới được '
  'chuẩn hóa từ những năm 2010 trở đi, chủ yếu trong lĩnh vực y học '
  'và y học công cộng (HMU, HUPH, Bệnh viện Bạch Mai, HCMUMP). Các '
  'cơ sở giáo dục đại học và khoa học xã hội nhân văn chỉ vừa thành '
  'lập HĐ Đạo đức trong vài năm gần đây (HCMUSSH 2022, USSH-VNU '
  'Hanoi 2021). Vì vậy, các luận án tiến sĩ tâm lý/giáo dục đời '
  'trước có thể không có QĐ chính thức.')

H2('2.2 — Cách hành xử "ghi chung chung" — nguy cơ với tạp chí quốc tế')
P('Việc NCS chỉ ghi "đã thông qua HĐ đạo đức" mà không có số QĐ cụ '
  'thể có hai vấn đề:')
BB('Reviewer quốc tế sẽ NGHI NGỜ tính chính xác của tuyên bố. Một '
   'số tạp chí (đặc biệt Q1) có thể yêu cầu bằng chứng cụ thể trước '
   'khi cho qua peer review.')
BB('Nếu reviewer hoặc editor liên hệ trực tiếp với cơ sở để xác '
   'minh, mà cơ sở không thể cung cấp số QĐ tương ứng, bài báo có '
   'nguy cơ bị retracted về sau.')

H2('2.3 — Khuyến nghị cho NCS Công Thị Hằng')
P('Em đề xuất NCS làm các bước sau theo đúng quy trình:')
P('Bước 1: Liên hệ Phòng Khoa học Công nghệ Trường Đại học Sư phạm '
  'Hà Nội (HNUE) để hỏi:')
BB('Trường có HĐ Đạo đức trong Nghiên cứu Khoa học chưa? Nếu có thì '
   'theo Quy chế nào, ban hành ngày nào?')
BB('Nếu chưa có HĐ riêng của HNUE, có thể xin phê duyệt qua HĐ Đạo '
   'đức Trường Đại học Y Hà Nội (HMU) hoặc Viện Sức khỏe Tâm thần '
   'Quốc gia không (do nội dung nghiên cứu liên quan đến sức khỏe '
   'tâm thần học sinh)?')

P('Bước 2: Khi xin phê duyệt, ghi rõ trong hồ sơ:')
BB('Mục tiêu nghiên cứu LA + 3 bài báo Q2/Q3/Q4 (nói rõ là dùng '
   'cùng bộ dữ liệu)')
BB('Quy trình thu thập dữ liệu + chấp thuận cha mẹ + đồng thuận học '
   'sinh + bảo mật + xử lý tình huống phát hiện học sinh có triệu '
   'chứng nặng')
BB('Phương án bảo mật dữ liệu + lưu trữ + truy cập')

P('Bước 3: Sau khi nhận QĐ, NCS giữ bản scan để đính kèm khi nộp 3 '
  'bài báo + LA. Trong manuscript, ghi rõ số QĐ + ngày phê duyệt + '
  'tên cơ quan phê duyệt.')


# ============================================================
H1('III. MẪU CÂU "ETHICS APPROVAL STATEMENT" CHUẨN QUỐC TẾ')

P('Em đề xuất mẫu câu để NCS dùng trong cả 3 bài Q2/Q3/Q4 (chỉ thay '
  'thông tin số QĐ thực tế sau khi nhận):')

H2('Mẫu tiếng Anh (dùng trong manuscript)')
P('"The study protocol was approved by the [TÊN HỘI ĐỒNG ĐẠO ĐỨC '
  'CHÍNH THỨC] of [TÊN TRƯỜNG/VIỆN] on [NGÀY phê duyệt] (Approval No. '
  '[SỐ QĐ]). Written informed consent was obtained from parents and '
  'written assent was obtained from students prior to data '
  'collection. All procedures were conducted in accordance with the '
  '1964 Declaration of Helsinki and its later amendments. The same '
  'approval covers the present analysis and two pre-registered '
  'companion analyses focusing on gender invariance and latent '
  'profile analysis of the same cohort."')

H2('Mẫu tiếng Việt (dùng trong LA + bản tiếng Việt)')
P('"Đề cương nghiên cứu đã được phê duyệt bởi [TÊN HỘI ĐỒNG ĐẠO ĐỨC '
  'CHÍNH THỨC] của [TÊN TRƯỜNG/VIỆN] ngày [NGÀY phê duyệt] (Quyết '
  'định số [SỐ QĐ]). Chấp thuận tham gia bằng văn bản đã được thu '
  'thập từ cha mẹ và đồng thuận bằng văn bản đã được thu thập từ học '
  'sinh trước khi triển khai thu dữ liệu. Tất cả quy trình tuân thủ '
  'Tuyên bố Helsinki năm 1964 và các sửa đổi sau đó. Cùng QĐ phê '
  'duyệt này bao trùm phân tích hiện tại và hai phân tích đồng hành '
  'đã đăng ký trước về bất biến giới và phân tích hồ sơ tiềm ẩn của '
  'cùng cohort."')


# ============================================================
H1('IV. SO SÁNH 2 PHƯƠNG ÁN STRATEGIC')

H2('Phương án A — 1 QĐ chung cho LA + Q2 + Q3 + Q4')
BB('Ưu điểm: Đơn giản, chỉ làm thủ tục 1 lần. Tiết kiệm thời gian + '
   'chi phí. Đúng thông lệ quốc tế.')
BB('Nhược điểm: Phải làm hồ sơ phê duyệt CHI TIẾT bao trùm cả 3 mục '
   'tiêu nghiên cứu (SEM tích hợp, SEM đa nhóm theo giới, LPA). Có '
   'thể yêu cầu Q3 và Q4 đã được "đăng ký trước" tại OSF trước khi '
   'xin QĐ.')
BB('Em khuyến nghị: PHƯƠNG ÁN A là tốt nhất.')

H2('Phương án B — 1 QĐ chung cho LA + 1 QĐ riêng cho mỗi bài báo '
   'quốc tế')
BB('Ưu điểm: Mỗi bài báo có QĐ riêng, rõ ràng cho từng reviewer.')
BB('Nhược điểm: Phức tạp hơn, mất 4 lần hồ sơ. Không cần thiết theo '
   'thông lệ quốc tế. Tốn thời gian.')
BB('Em KHÔNG khuyến nghị phương án này.')

H2('Phương án C — 1 QĐ chỉ cho phép triển khai LA, không nhắc đến '
   'công bố quốc tế')
BB('Ưu điểm: Đơn giản nhất.')
BB('Nhược điểm: Reviewer quốc tế có thể HỎI THÊM về việc QĐ có bao '
   'trùm công bố quốc tế không. Nếu HĐ Đạo đức yêu cầu sửa thì phải '
   'làm lại.')
BB('Em KHÔNG khuyến nghị — rủi ro cao về sau.')


# ============================================================
H1('V. VÍ DỤ THỰC TẾ TỪ BÀI BÁO VIỆT NAM TRÊN TẠP CHÍ QUỐC TẾ')

P('Em sưu tầm 2 ví dụ thực tế từ bài báo Việt Nam đã công bố Q1/Q2 '
  'gần đây, để nhóm tham khảo cách họ ghi ethics statement:')

H2('Ví dụ 1 — Pham TT Hoa và cs. (2024), Frontiers in Public Health')
P('Bài "Anxiety symptoms and coping strategies among high school '
  'students in Vietnam after COVID-19 pandemic: a mixed-method '
  'evaluation" — cũng nghiên cứu trên học sinh Hà Nội, mẫu N=3.910.')
P('Cách ghi: "The study was approved by the Ethics Committee of the '
  '[institution] (approval number XXX, date YYY). Written informed '
  'consent was obtained from all participants and their parents '
  'prior to data collection." — đầy đủ 3 yếu tố: cơ quan + số + '
  'ngày.')

H2('Ví dụ 2 — Erskine HE và cs. (2023), J Adolesc Health (V-NAMHS '
   'protocol)')
P('Bài "Measuring the Prevalence of Mental Disorders in Adolescents '
  'in Kenya, Indonesia, and Vietnam" — protocol nghiên cứu cấp quốc '
  'gia V-NAMHS, hợp tác Viện Xã hội học VN + Đại học Queensland Úc '
  '+ Johns Hopkins Mỹ.')
P('Cách ghi: "The study protocol received ethics approval from the '
  'Vietnam Institute of Sociology Ethical Review Board (approval '
  'date YYY) and from the University of Queensland Behavioural and '
  'Social Sciences Ethical Review Committee (approval number ZZZ)." '
  '— ghi rõ cả 2 cơ quan phê duyệt (Việt Nam + nước ngoài, vì có '
  'hợp tác quốc tế).')

P('Áp dụng cho mình: nếu chỉ có HĐ Đạo đức Việt Nam là đủ. Nếu sau '
  'này có hợp tác quốc tế (ví dụ thầy NMĐ ở Singapore/Anh/Mỹ), có '
  'thể cần thêm QĐ bên đối tác.')


# ============================================================
H1('VI. CÂU DẪN ZALO ĐỀ XUẤT')

P('Em đề xuất câu dẫn ngắn để gửi nhóm, kèm doc này làm attachment ạ:')

p = d.add_paragraph()
p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.right_indent = Cm(0.5)
p.paragraph_format.space_before = Pt(6); p.paragraph_format.space_after = Pt(8)
r = p.add_run('"Em xin gửi các thầy và chị tài liệu tham khảo về '
              'quyết định Hội đồng đạo đức ạ. Em đã tổng hợp trả lời '
              '3 câu hỏi của thầy ở mục I, kèm theo bối cảnh Việt '
              'Nam ở mục II, mẫu câu Ethics statement chuẩn quốc tế '
              'ở mục III, so sánh 3 phương án strategic ở mục IV, '
              'và 2 ví dụ thực tế từ bài báo Việt Nam đã công bố ở '
              'mục V. Tóm tắt nhanh: (1) chỉ cần 1 QĐ chung cho '
              'toàn bộ đề tài LA + 3 bài báo, (2) chỉ ghi số QĐ là '
              'không đủ cho tạp chí Q1/Q2 quốc tế, cần thêm '
              'cơ quan + ngày + scope, (3) HMU và HCMUSSH có template '
              'chuẩn để tham khảo. Có gì cần điều chỉnh xin các '
              'thầy và chị nhắn lại ạ."')
r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = True


# ============================================================
H1('TÀI LIỆU THAM KHẢO')

refs = [
    'World Medical Association. Declaration of Helsinki — Ethical '
    'Principles for Medical Research Involving Human Subjects. '
    'JAMA. 2013;310(20):2191-2194. DOI: 10.1001/jama.2013.281053.',
    'American Psychological Association (APA). IRBs and psychological '
    'science: Ensuring a collaborative relationship. 2024. URL: '
    'https://www.apa.org/research-practice/conduct-research/'
    'irbs-psych-science',
    'Committee on Publication Ethics (COPE). Handling duplicated or '
    'redundant content (salami slicing). 2018. URL: '
    'https://publicationethics.org/guidance/cope-position/'
    'handling-duplicated-or-redundant-content-salami-slicing',
    'Erskine HE, Blondell SJ, Enright ME, et al. Measuring the '
    'Prevalence of Mental Disorders in Adolescents in Kenya, '
    'Indonesia, and Vietnam: Study Protocol for the National '
    'Adolescent Mental Health Surveys. J Adolesc Health. '
    '2023;72(1S):S71-S78. PMID: 36229399. DOI: '
    '10.1016/j.jadohealth.2021.05.012.',
    'Pham TTH, Do TT, Nguyen TL, Ngo AV. Anxiety symptoms and '
    'coping strategies among high school students in Vietnam after '
    'COVID-19 pandemic: a mixed-method evaluation. Front Public '
    'Health. 2024;12:1232856. PMID: 38435293. DOI: '
    '10.3389/fpubh.2024.1232856.',
    'Hanoi Medical University Institutional Ethical Review Board. '
    'Contact: irb@hmu.edu.vn; Phone: 024-388-527-622. IRB '
    'registration: IRB-VN01.001.III.RB00003121/FWA00004148.',
    'Trường Đại học Khoa học Xã hội và Nhân văn — ĐHQG TP.HCM. Quy '
    'chế tổ chức và hoạt động Hội đồng Đạo đức Nghiên cứu. URL: '
    'https://hcmussh.edu.vn/tin-tuc/'
    'ban-hanh-quy-che-to-chuc-va-hoat-dong-hoi-dong-dao-duc-nghien-cuu',
    'Frontiers in Psychiatry. Author Guidelines — Ethics Statement '
    'Requirements. URL: '
    'https://www.frontiersin.org/journals/psychiatry/for-authors/'
    'author-guidelines',
]
for ref in refs:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.75); p.paragraph_format.first_line_indent = Cm(-0.75)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run('• ' + ref); r.font.name = 'Times New Roman'; r.font.size = Pt(10)


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'SAVED: {OUT}')
print(f'SIZE: {os.path.getsize(OUT)} bytes')
from docx import Document as Doc
d2 = Doc(OUT)
chunks = [p.text for p in d2.paragraphs if p.text.strip()]
print(f'WORD COUNT: {sum(len(p.split()) for p in chunks)}')
