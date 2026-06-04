# -*- coding: utf-8 -*-
"""Rebuild bang TOC table tu cac heading thuc te trong doc.
- Xoa cac row cu (TOC outdated)
- Them tat ca H1-H3 (skip H4 de tranh qua dai)
- Page so chinh xac tu Word COM render
- Giu cau truc bang (2 cot, font, format)
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import win32com.client, pythoncom
from docx import Document
from docx.shared import Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from copy import deepcopy

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

# Backup
BACKUP = FILE.replace('.docx', '_BEFORE_TOC_REBUILD2.docx')
shutil.copy(FILE, BACKUP)
print(f"Backup: {BACKUP}")


# ============================================================
# 1. Extract headings + pages via Word COM
# ============================================================
print("\n[1] Extract headings...")
TEMP = os.path.join(ROOT, 'Luận án TS', '_temp_rebuild.docx')
shutil.copy(FILE, TEMP)
pythoncom.CoInitialize()
word = None
headings = []
try:
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False; word.DisplayAlerts = False
    doc = word.Documents.Open(os.path.abspath(TEMP), ReadOnly=False)
    doc.Fields.Update(); doc.Repaginate()
    for level in [1, 2, 3, 4, 5]:
        for style_name in [f'Heading {level}', f'Tiêu đề {level}']:
            try:
                rng = doc.Content
                find = rng.Find
                find.ClearFormatting(); find.Style = doc.Styles(style_name)
                find.Text = ''; find.Forward = True; find.Wrap = 0; find.Format = True
                while find.Execute():
                    para = rng.Paragraphs.First
                    text = para.Range.Text.strip()
                    # Remove ALL control characters (XML-incompatible)
                    text = ''.join(ch for ch in text if ord(ch) >= 32 or ch in '\t\n')
                    text = text.strip()
                    if para.Range.Tables.Count > 0: rng.Collapse(0); continue
                    if not text or len(text) > 250: rng.Collapse(0); continue
                    page = para.Range.Information(3)
                    pos = para.Range.Start
                    headings.append({'level': level, 'text': text, 'page': page, 'pos': pos})
                    rng.Collapse(0)
            except: pass
    doc.Close(SaveChanges=False)
finally:
    if word:
        try: word.Quit()
        except: pass
    pythoncom.CoUninitialize()
    if os.path.exists(TEMP):
        try: os.remove(TEMP)
        except: pass

# Dedupe + sort by document position
seen = set(); unique = []
for h in headings:
    k = (h['text'], h['page'])
    if k not in seen:
        seen.add(k); unique.append(h)
headings = sorted(unique, key=lambda h: h['pos'])
print(f"  {len(headings)} unique headings")


# ============================================================
# 2. Filter: keep H1-H3, skip cover-type (DANH MỤC, etc.)
# ============================================================
print("\n[2] Filter for TOC (H1-H3, skip cover)...")
SKIP_TEXTS = {'MỤC LỤC', 'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
              'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
              'LỜI CAM ĐOAN', 'LỜI CẢM ƠN'}
filtered = []
for h in headings:
    if h['level'] > 3:
        continue
    if h['text'].upper().strip() in SKIP_TEXTS:
        continue
    filtered.append(h)
print(f"  Filtered: {len(filtered)} headings for TOC")


# ============================================================
# 3. Compose hierarchical display text + page list
# ============================================================
print("\n[3] Compose TOC entries...")
# Special handling: combine CHƯƠNG X (separate paragraph) with chapter title (next para)
# E.g., "CHƯƠNG 1" + "CƠ SỞ LÝ LUẬN..." → "CHƯƠNG 1: CƠ SỞ LÝ LUẬN..."
toc_entries = []  # list of (display_text, page, level)
i = 0
while i < len(filtered):
    h = filtered[i]
    text = h['text'].strip()
    if re.match(r'^CHƯƠNG\s+\d+$', text):
        # Look at next heading
        if i + 1 < len(filtered) and filtered[i+1]['level'] == 1 and filtered[i+1]['page'] - h['page'] < 3:
            # Combine "CHƯƠNG X" + title
            combined = f"{text}: {filtered[i+1]['text'].strip()}"
            toc_entries.append((combined, h['page'], 1))
            i += 2
            continue
    toc_entries.append((text, h['page'], h['level']))
    i += 1

print(f"  {len(toc_entries)} TOC entries composed")


# ============================================================
# 4. Find existing TOC table + capture style from first row
# ============================================================
d = Document(FILE)
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb
        break

if toc_table is None:
    print("KHONG TIM THAY TOC TABLE")
    sys.exit(1)

# Capture style from existing first row
first_row = toc_table.rows[0]
sample_cell0 = first_row.cells[0]
sample_cell1 = first_row.cells[1]
# Get font info from first paragraph first run
def get_run_style(cell):
    if cell.paragraphs and cell.paragraphs[0].runs:
        r = cell.paragraphs[0].runs[0]
        return {
            'font_name': r.font.name or 'Times New Roman',
            'font_size': r.font.size or Pt(13),
            'bold': r.bold,
        }
    return {'font_name': 'Times New Roman', 'font_size': Pt(13), 'bold': None}

style0 = get_run_style(sample_cell0)
style1 = get_run_style(sample_cell1)
print(f"  Cell0 style: {style0}")
print(f"  Cell1 style: {style1}")


# ============================================================
# 5. Rebuild: clear all rows except first row template, then re-add
# ============================================================
print("\n[5] Rebuilding TOC table...")
n_rows = len(toc_table.rows)
print(f"  Original rows: {n_rows}")

# Delete all rows
# Direct XML manipulation
tbl = toc_table._element
# Find all <w:tr> children
rows_xml = tbl.findall(qn('w:tr'))
for tr in rows_xml:
    tbl.remove(tr)
print(f"  Cleared all rows")

# Now add new rows
def add_row_to_table(table, name, page, level):
    """Add a new row with given name + page. Style based on level."""
    row = table.add_row()
    # Cell 0: name
    cell0 = row.cells[0]
    cell0.paragraphs[0].text = ''
    p0 = cell0.paragraphs[0]
    r0 = p0.add_run(name)
    r0.font.name = 'Times New Roman'
    if level == 1:
        r0.font.size = Pt(13)
        r0.bold = True
    elif level == 2:
        r0.font.size = Pt(13)
        r0.bold = True
    elif level == 3:
        r0.font.size = Pt(13)
        r0.bold = False
    # Cell 1: page
    cell1 = row.cells[1]
    cell1.paragraphs[0].text = ''
    p1 = cell1.paragraphs[0]
    r1 = p1.add_run(str(page))
    r1.font.name = 'Times New Roman'
    r1.font.size = Pt(13)
    if level == 1:
        r1.bold = True
    return row

for name, page, level in toc_entries:
    # Clean any remaining control chars
    name = ''.join(ch for ch in str(name) if ord(ch) >= 32 or ch in '\t\n').strip()
    if not name:
        continue
    # Indent for visual hierarchy
    indent = '   ' * (level - 1) if level > 1 else ''
    display = indent + name
    add_row_to_table(toc_table, display, page, level)

print(f"  Added {len(toc_entries)} new rows")

d.save(FILE)
print(f"\nSaved: {FILE}")
print(f"Size: {os.path.getsize(FILE)//1024} KB")

# Verify
d2 = Document(FILE)
new_toc = None
for tb in d2.tables:
    if len(tb.rows) > 0:
        first = tb.rows[0].cells[0].text.strip()
        if first.startswith('MỞ ĐẦU') or 'MỞ ĐẦU' in first:
            new_toc = tb
            break
if new_toc:
    print(f"\nNew TOC: {len(new_toc.rows)} rows")
    print("First 20 rows:")
    for i, row in enumerate(new_toc.rows[:20]):
        print(f"  Row {i}: {row.cells[0].text.strip()[:70]:<70} | {row.cells[1].text.strip()}")
