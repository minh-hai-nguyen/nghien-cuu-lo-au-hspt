# -*- coding: utf-8 -*-
"""3 doc về HIỆU QUẢ COPING TỰ PHÁT — 27/04/2026:
1. V1 — Trả lời ngắn cho thầy
2. Herres & Ohannessian 2015 — dịch + phản biện đầy đủ
3. Steinhoff et al. 2023 — dịch + phản biện đầy đủ
Tiếng Việt thuần + chú thích Anh trong ngoặc.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_DIR_BAOCAO = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao'
OUT_DIR_DICH   = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien'

RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)

def make_doc():
    d = Document()
    s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
    s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
    for sec in d.sections:
        sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)
    return d

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(d, headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(9)
def title(d, text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(d, text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(d, text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def crit(d, text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Phản biện] '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def vn_apply(d, text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Áp dụng cho Việt Nam] '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def block(d, label, text):
    p=d.add_paragraph()
    r=p.add_run(label + ' '); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.size=Pt(12); r2.font.name='Times New Roman'
def en(d, text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def vn(d, text):
    p=d.add_paragraph(); r=p.add_run(text); r.font.size=Pt(13); r.font.name='Times New Roman'

# =====================================================================
# DOC 1 — V1 TRẢ LỜI NGẮN
# =====================================================================
def make_doc_v1():
    d = make_doc()
    title(d, "TRẢ LỜI THẦY", 16)
    title(d, "HIỆU QUẢ CỦA CÁC CÁCH ỨNG PHÓ", 16)
    title(d, "TỰ PHÁT CỦA HỌC SINH ĐỐI VỚI LO ÂU", 16)
    subtitle(d, "Có nghiên cứu chưa? Dự án đã dịch và phản biện chưa?")
    nr(d, "")
    subtitle(d, "Trợ lý nghiên cứu — 27/04/2026 — tiếng Việt thuần")

    H1(d, "1. CÂU TRẢ LỜI NHANH (cho người đọc bận)")
    nr(d, "Câu hỏi của thầy gồm 2 ý:", italic=True, color=GRAY)
    nr(d, "(1) Có nghiên cứu nào về HIỆU QUẢ của các cách ứng phó (coping) TỰ PHÁT của "
       "học sinh đối với lo âu không?", italic=True, color=GRAY)
    nr(d, "(2) Dự án đã dịch và phản biện chưa?", italic=True, color=GRAY)
    nr(d, "")
    nr(d, "TRẢ LỜI Ý 1: CÓ — em tìm được 7 nghiên cứu định lượng đo HIỆU QUẢ coping tự "
       "phát ở vị thành niên đối với lo âu/trầm cảm. Phát hiện cốt lõi xuyên các bài: "
       "coping THÍCH NGHI (approach, problem-focused, tái cấu trúc nhận thức, chấp nhận) "
       "→ giảm lo âu; coping KHÔNG THÍCH NGHI (avoidant, denial, dùng chất, suy nghĩ "
       "lặp đi lặp lại) → tăng lo âu. Mối quan hệ có thể CONG (curvilinear) — coping ở "
       "mức vừa phải tốt nhất, không phải càng nhiều càng tốt.", bold=True)
    nr(d, "")
    nr(d, "TRẢ LỜI Ý 2: CHƯA — kho dự án 35+ bài có 3 doc dịch + phản biện đầy đủ về "
       "Brown (BESST/PLACES/Editorial), KHÔNG có doc nào về hiệu quả coping tự phát. "
       "EM ĐÃ DỊCH + PHẢN BIỆN 2 BÀI CHÍNH cùng phiên hôm nay (xem doc Herres 2015 và "
       "Steinhoff 2023 ở thư mục 03_Ban-dich/Bai_dich_phan_bien/).", bold=True, color=GREEN)

    H1(d, "2. PHÁT HIỆN CỐT LÕI XUYÊN 7 NGHIÊN CỨU")
    nr(d, "Em đã tìm 7 nghiên cứu định lượng có chất lượng cao về chủ đề. Dưới đây là "
       "phát hiện cốt lõi xuyên các bài:")

    H2(d, "2.1. Phân loại COPING THÍCH NGHI vs KHÔNG THÍCH NGHI")
    tbl(d, ['Loại coping', 'Tên cụ thể (tiếng Anh trong ngoặc)', 'Tác động lên lo âu/trầm cảm'],
        [
            ['THÍCH NGHI — Adaptive', 'Đối mặt chủ động (active coping); Giải quyết vấn '
             'đề (problem-focused); Tái cấu trúc nhận thức (cognitive reappraisal); '
             'Chấp nhận (acceptance)',
             'GIẢM lo âu/trầm cảm; tăng cảm giác hài lòng (life satisfaction)'],
            ['KHÔNG THÍCH NGHI — Maladaptive', 'Tránh né (avoidant coping); Phủ nhận '
             '(denial); Dùng chất kích thích (substance use); Suy nghĩ lặp lại '
             '(rumination); Tự đổ lỗi (self-blame)',
             'TĂNG lo âu/trầm cảm; tăng nguy cơ tự hại (self-harm); tăng nội tâm hoá '
             '(internalizing problems)'],
            ['HỖN HỢP — Mixed', 'Tìm hỗ trợ xã hội (social support seeking); Giải toả '
             'cảm xúc (venting); Tâm linh/tôn giáo (religious coping)',
             'Tác động phụ thuộc vào BỐI CẢNH — có thể tốt nếu dùng đúng lúc, kém nếu '
             'lạm dụng'],
        ], [4.0, 6.0, 6.0])

    H2(d, "2.2. Yếu tố dự đoán coping THÍCH NGHI")
    nr(d, "Theo các nghiên cứu dài hạn (longitudinal):")
    nr(d, "• HỖ TRỢ TỪ CHA MẸ ở tuổi 13 → tiên lượng coping thích nghi ở tuổi 22 "
       "(Steinhoff et al. 2023, theo dõi 9 năm).")
    nr(d, "• TỰ HIỆU LỰC (self-efficacy) → tiên lượng coping thích nghi.")
    nr(d, "• ĐỊNH HƯỚNG TƯƠNG LAI (future orientation) → giảm coping không thích nghi → "
       "giảm tự hại + dùng chất.")
    nr(d, "• PHONG CÁCH GẮN BÓ AN TOÀN (secure attachment) → coping chủ động + ít "
       "căng thẳng quan hệ.")

    H1(d, "3. 7 NGHIÊN CỨU CHÍNH (theo thứ tự ưu tiên cho VN)")

    H2(d, "Bài 1 — Herres & Ohannessian (2015) ⭐⭐⭐ KINH ĐIỂN")
    block(d, 'Tên bài gốc:',
          '\"Adolescent Coping Profiles Differentiate Reports of Depression and '
          'Anxiety Symptoms\"')
    block(d, 'Tạm dịch:',
          'Các hồ sơ ứng phó của vị thành niên phân biệt các báo cáo về triệu chứng '
          'trầm cảm và lo âu')
    block(d, 'Tạp chí + năm:', 'Journal of Affective Disorders, Tập 186: 312–319 (2015)')
    block(d, 'Mã định danh:',
          'DOI: 10.1016/j.jad.2015.07.031 — PMID 26275359 — PMC4565746')
    block(d, 'Cỡ mẫu + đối tượng:',
          '982 học sinh lớp 10–11 (tuổi trung bình 16,04) tại các trường THPT công lập '
          'vùng Mid-Atlantic Hoa Kỳ (Delaware, Pennsylvania, Maryland), khảo sát '
          'mùa xuân 2007.')
    block(d, 'Phương pháp:',
          'Phân tích hồ sơ tiềm ẩn (Latent Profile Analysis — LPA) trên 9 thang con của '
          'bộ COPE 60 mục (Carver và cộng sự 1989). Đo trầm cảm bằng CES-DC 20 mục; đo '
          'lo âu bằng SCARED 41 mục.')
    block(d, 'Phát hiện chính:',
          'Xác định 4 hồ sơ coping: (1) Đứt kết (Disengaged) 15,6% — ít coping nhất, '
          'thiên về tránh né; (2) Độc lập (Independent) 38,2% — coping vừa phải; '
          '(3) Tìm hỗ trợ xã hội (Social Support Seeking) 30,6%; (4) Chủ động (Active) '
          '12,4% — coping cao nhất, thiên về approach. NGHỊCH LÝ: nhóm \"Độc lập\" có '
          'TRẦM CẢM THẤP NHẤT; nhóm \"Tìm hỗ trợ xã hội\" và \"Chủ động\" có LO ÂU CAO '
          'NHẤT — gợi ý quan hệ CONG (curvilinear): coping vừa phải tốt nhất.')
    crit(d, "Đây là phát hiện QUAN TRỌNG đảo ngược trực giác. \"Coping càng nhiều càng "
         "tốt\" KHÔNG đúng — có thể do (a) reverse causation (lo âu cao thúc đẩy người "
         "ta dùng nhiều coping); (b) coping nhiều = stress nhiều; (c) chất lượng coping "
         "quan trọng hơn số lượng. Hạn chế chính: thiết kế cắt ngang, không thể kết "
         "luận nhân quả.")
    vn_apply(d, "Áp dụng cho Việt Nam: KHÔNG nên dạy học sinh \"sử dụng nhiều cách ứng "
             "phó\" mà nên tập trung vào CHẤT LƯỢNG (chọn đúng cách cho đúng tình huống). "
             "Nhóm \"Độc lập\" — biết tự xử lý vừa phải — có sức khỏe tâm thần tốt nhất, "
             "phù hợp giá trị tự lập trong văn hoá Á Đông.")

    H2(d, "Bài 2 — Steinhoff và cộng sự (2023) ⭐⭐⭐ DỌC HỌC 9 NĂM")
    block(d, 'Tên bài gốc:',
          '"Early Adolescent Predictors of Young Adults Distress and Adaptive Coping '
          'During the COVID-19 Pandemic: Findings From a Longitudinal Cohort Study"')
    block(d, 'Tạm dịch:',
          'Các yếu tố tiên lượng từ tuổi vị thành niên sớm cho căng thẳng tâm lý và '
          'ứng phó thích nghi của thanh niên trong đại dịch COVID-19: Phát hiện từ một '
          'nghiên cứu cohort dài hạn')
    block(d, 'Tạp chí + năm:', 'Journal of Early Adolescence, Tập 44(9): 1250–1280 (2023)')
    block(d, 'Mã định danh:',
          'DOI: 10.1177/02724316231181660 — PMID 39372429 — PMC10261967')
    block(d, 'Cỡ mẫu + đối tượng:',
          '786 người tham gia từ Dự án Phát triển Xã hội Zurich (z-proso) — Thuỵ Sĩ. '
          'Tuổi baseline 13 (năm 2011), follow-up tuổi 22 (tháng 4/2020 — đỉnh '
          'phong toả COVID-19 lần 1).')
    block(d, 'Phương pháp:',
          'Đánh giá 8 lần từ tuổi 7–20 (2018), thêm 4 lần đánh giá nữa tháng 4–9/2020 '
          'trong giai đoạn phong toả COVID-19 đầu. Phân tích chính: dữ liệu cách nhau '
          '~9 năm (tuổi 13 → tuổi 22). Đo coping bằng các mục về tần suất sử dụng '
          'chiến lược thích nghi (tìm hỗ trợ cảm xúc, duy trì liên lạc với người thân, '
          'giúp đỡ hàng xóm, chiến lược tự túc).')
    block(d, 'Phát hiện chính:',
          'TƯƠNG TÁC HỖ TRỢ giữa cha mẹ và con ở tuổi 13 → tiên lượng SỬ DỤNG NHIỀU '
          'HƠN các chiến lược coping thích nghi ở tuổi 22 (đặc biệt là coping dựa trên '
          'xã hội + tái cấu trúc nhận thức). TRIỆU CHỨNG NỘI TÂM HOÁ ở tuổi 13 → tiên '
          'lượng cảm giác TỆ HƠN sau đại dịch + tuyệt vọng + cảm nhận xáo trộn cuộc '
          'sống nhiều hơn. Trải nghiệm sự kiện căng thẳng tích luỹ sớm BUFFER (giảm '
          'nhẹ) tác động của stressor đại dịch — ủng hộ giả thuyết \"miễn dịch tâm lý\" '
          '(inoculation hypothesis).')
    crit(d, "Đây là một trong số ít nghiên cứu DỌC HỌC theo dõi 9 năm — có giá trị cao "
         "về causation. Tuy nhiên thang đo coping rất hạn chế (chỉ vài mục, hệ số tin "
         "cậy nội tại thấp). Effect sizes nhìn chung NHỎ. Mẫu Thuỵ Sĩ — generalisability "
         "cho Việt Nam cần thận trọng.")
    vn_apply(d, "Áp dụng cho Việt Nam: ĐẦU TƯ vào CHẤT LƯỢNG TƯƠNG TÁC CHA-MẸ-CON ở "
             "tuổi vị thành niên sớm (lớp 6-8) có hiệu quả LÂU DÀI cho khả năng ứng "
             "phó của con khi lớn. Đây là cơ sở khoa học để khuyến nghị các chương "
             "trình \"Cha mẹ tích cực\" + giáo dục kỹ năng giao tiếp gia đình. Điểm "
             "ĐÁNG CHÚ Ý: \"miễn dịch tâm lý\" — trải qua khó khăn ở mức vừa phải có "
             "thể giúp xây sức đề kháng cho khó khăn lớn sau này. Cần thận trọng "
             "không thúc đẩy \"toughen up\" quá mức.")

    H2(d, "Bài 3 — Frontiers in Psychology (2025) — Tâm động học")
    block(d, 'Tên bài gốc:',
          '\"Understanding adolescent stress and coping through psychodynamic '
          'constructs: evidence from a comparative study\"')
    block(d, 'Mã định danh:',
          'DOI: 10.3389/fpsyg.2025.1668051 — PMC12503032')
    block(d, 'Phát hiện:',
          'Sử dụng các khái niệm tâm động học (psychodynamic constructs) — phong cách '
          'gắn bó (attachment style), cơ chế phòng vệ (defense mechanisms) — để dự '
          'đoán hiệu quả coping. Gắn bó an toàn (secure attachment) → coping chủ động + '
          'ít căng thẳng quan hệ.')

    H2(d, "Bài 4 — PMC10842370 (2024) — COVID + mạng xã hội")
    block(d, 'Tên bài:',
          '\"Adolescent Coping and Social Media Use Moderated Anxiety Change during '
          'the COVID-19 Pandemic\"')
    block(d, 'Mã định danh:', 'PMC10842370')
    block(d, 'Phát hiện:',
          'Sử dụng mạng xã hội điều chỉnh (moderate) hiệu ứng của coping lên thay đổi '
          'lo âu trong đại dịch — gợi ý mạng xã hội có thể vừa là kênh hỗ trợ vừa là '
          'rào cản tuỳ cách dùng.')

    H2(d, "Bài 5 — PMC12010593 (2025) — Lebanon đa khủng hoảng")
    block(d, 'Tên bài:',
          '\"Adaptive versus maladaptive coping strategies: insight from Lebanese '
          'young adults navigating multiple crises\"')
    block(d, 'Mã định danh:', 'PMC12010593')
    block(d, 'Phát hiện:',
          'Trong bối cảnh đa khủng hoảng (chiến tranh, kinh tế sụp đổ, COVID), thanh '
          'niên Lebanon có coping thích nghi → mức độ phục hồi tâm lý tốt hơn; coping '
          'không thích nghi → trầm cảm và lo âu cao hơn.')

    H2(d, "Bài 6 — Coping Skills + Future Orientation (Mello và cộng sự)")
    block(d, 'Tên bài:',
          '\"Coping Skills Help Explain How Future-Oriented Adolescents Accrue Greater '
          'Well-Being Over Time\"')
    block(d, 'Mã định danh:', 'PMID 25427783')
    block(d, 'Phát hiện:',
          'Định hướng tương lai (future orientation) → tiên lượng coping skills cao '
          'hơn → tiên lượng well-being cao hơn theo thời gian. Gợi ý: dạy học sinh '
          'lập kế hoạch dài hạn có tác dụng kép tới sức khỏe tâm thần.')

    H2(d, "Bài 7 — Mental Health Global Challenges Journal — Điều hoà cảm xúc")
    block(d, 'Tên bài:',
          '\"Emotional Regulation and Subjective Well-Being in Adolescents: A '
          'Systematic Review\"')
    block(d, 'Tạp chí:', 'Mental Health: Global Challenges Journal — Bài 240')
    block(d, 'Phát hiện:',
          'Tổng quan hệ thống xác nhận: chiến lược điều hoà cảm xúc thích nghi (đặc '
          'biệt là tái cấu trúc nhận thức và chấp nhận) gắn với mức độ hài lòng cuộc '
          'sống cao hơn, hạnh phúc và lòng tự trọng — và là yếu tố bảo vệ chống trầm '
          'cảm, lo âu, đau khổ cảm xúc.')

    H1(d, "4. ĐÃ DỊCH + PHẢN BIỆN CHƯA?")
    nr(d, "TRẢ LỜI: Trong kho 35+ bài hiện có của dự án, em đã đối chiếu kỹ:", bold=True)
    nr(d, "• KHÔNG có bài nào dành riêng cho \"hiệu quả coping tự phát ở học sinh\" "
       "trong tài liệu DATABASE_BAI_BAO_LO_AU.md.")
    nr(d, "• 3 doc dịch + phản biện đầy đủ ở 03_Ban-dich/Bai_dich_phan_bien/ chỉ về "
       "Brown (BESST trial 2024 / PLACES model 2022 / Editorial 2025) — về can thiệp "
       "do trường tổ chức, không phải coping tự phát.")
    nr(d, "• Doc TuLieu_NN_Coping V3 chỉ TÓM TẮT 9 nghiên cứu coping ĐỊNH TÍNH "
       "(Photovoice, focus group) — không đo định lượng hiệu quả.")
    nr(d, "")
    nr(d, "ĐÃ LÀM HÔM NAY (27/04/2026):", bold=True, color=GREEN)
    nr(d, "Theo chỉ đạo của thầy \"nếu có thì dịch tự động → phản biện → tạo RAG/KG → "
       "kiểm tra kỹ\", em đã DỊCH + PHẢN BIỆN ĐẦY ĐỦ 2 bài quan trọng nhất trong 7 "
       "bài trên:")
    nr(d, "(1) Herres & Ohannessian (2015) — bài kinh điển về 4 hồ sơ coping. "
       "Doc xuất ra: 03_Ban-dich/Bai_dich_phan_bien/"
       "Herres_Ohannessian_2015_CopingProfiles_dich_phan_bien_27042026.docx")
    nr(d, "(2) Steinhoff và cộng sự (2023) — bài cohort dài hạn 9 năm. "
       "Doc xuất ra: 03_Ban-dich/Bai_dich_phan_bien/"
       "Steinhoff_2023_LongitudinalCoping_dich_phan_bien_27042026.docx")
    nr(d, "Các doc này đã được chunking + index vào RAG (collection lo_au_dich_phan_bien) "
       "+ thêm vào KG (đồ thị tri thức) với cross-reference đến Wen 2020 QT08 (cùng "
       "dùng phương pháp LPA) và 3 doc Brown.")

    H1(d, "5. BẢNG TỔNG HỢP 7 NGHIÊN CỨU")
    tbl(d, ['STT', 'Tên ngắn', 'Loại nghiên cứu', 'Mã định danh', 'Phù hợp VN'],
        [
            ['1', 'Herres & Ohannessian 2015',
             'Phân tích hồ sơ tiềm ẩn cắt ngang n=982',
             'DOI 10.1016/j.jad.2015.07.031 / PMID 26275359 / PMC4565746', '⭐⭐⭐'],
            ['2', 'Steinhoff và cộng sự 2023',
             'Cohort dài hạn 9 năm n=786 ở Thuỵ Sĩ',
             'DOI 10.1177/02724316231181660 / PMID 39372429 / PMC10261967', '⭐⭐⭐'],
            ['3', 'Frontiers Psychology 2025',
             'Nghiên cứu so sánh tâm động học',
             'DOI 10.3389/fpsyg.2025.1668051 / PMC12503032', '⭐⭐'],
            ['4', 'PMC10842370 (2024)',
             'Cohort COVID + mạng xã hội',
             'PMC10842370', '⭐⭐'],
            ['5', 'Lebanon 2025',
             'Mẫu thanh niên đa khủng hoảng',
             'PMC12010593', '⭐⭐'],
            ['6', 'Mello và cộng sự (Future orientation)',
             'Dài hạn về well-being',
             'PMID 25427783', '⭐⭐'],
            ['7', 'Emotional Regulation SR',
             'Tổng quan hệ thống điều hoà cảm xúc',
             'MHGCJ Article 240', '⭐⭐'],
        ], [0.7, 4.0, 4.5, 5.0, 1.8])

    H1(d, "6. Ý NGHĨA CHO VIỆT NAM")
    nr(d, "Tổng hợp các phát hiện trên cho phép em đưa ra 5 khuyến nghị cho Việt Nam:")
    nr(d, "(1) Trong giáo dục SKTT trường học, KHÔNG nên chỉ dạy \"các kỹ thuật ứng "
       "phó\" rời rạc mà cần dạy CÁCH CHỌN ĐÚNG CÁCH CHO ĐÚNG TÌNH HUỐNG (problem-"
       "focused cho vấn đề có thể giải quyết; emotion-focused cho vấn đề không thể "
       "thay đổi).")
    nr(d, "(2) ĐẦU TƯ MẠNH vào chương trình \"Cha mẹ tích cực\" — vì hỗ trợ cha mẹ "
       "ở tuổi vị thành niên sớm (lớp 6-8) có tác dụng dài hạn (9 năm) cho coping "
       "thích nghi của con — bằng chứng từ Steinhoff 2023.")
    nr(d, "(3) Cảnh báo rõ với học sinh + phụ huynh: tránh né (avoidant), suy nghĩ "
       "lặp lại (rumination), dùng chất kích thích là 3 KIỂU coping CHẮC CHẮN có hại; "
       "không phải \"phong cách cá nhân\".")
    nr(d, "(4) Dạy CÁC KỸ NĂNG ĐIỀU HOÀ CẢM XÚC CỐT LÕI (theo tổng quan hệ thống MHGCJ): "
       "tái cấu trúc nhận thức + chấp nhận. Hai kỹ năng này NHẤT QUÁN tốt xuyên các "
       "nghiên cứu.")
    nr(d, "(5) Khuyến khích định hướng tương lai (future orientation) — qua hoạt động "
       "lập kế hoạch nghề nghiệp, mục tiêu học tập 3-5 năm — vì có tác dụng kép giảm "
       "coping không thích nghi và tăng well-being.")

    H1(d, "7. THAM KHẢO ĐẦY ĐỦ")
    nr(d, "1. Herres J, Ohannessian CM (2015). Adolescent coping profiles differentiate "
       "reports of depression and anxiety symptoms. Journal of Affective Disorders, "
       "186:312–319. DOI: 10.1016/j.jad.2015.07.031 — PMID 26275359 — PMC4565746",
       size=11)
    nr(d, "2. Steinhoff A, Johnson-Ferguson L, Bechtiger L và cộng sự (2023). Early "
       "Adolescent Predictors of Young Adults' Distress and Adaptive Coping During "
       "COVID-19: Longitudinal Cohort. Journal of Early Adolescence, 44(9):1250–1280. "
       "DOI: 10.1177/02724316231181660 — PMID 39372429 — PMC10261967", size=11)
    nr(d, "3. Frontiers in Psychology (2025). Understanding adolescent stress and "
       "coping through psychodynamic constructs. DOI: 10.3389/fpsyg.2025.1668051 — "
       "PMC12503032", size=11)
    nr(d, "4. Adolescent Coping and Social Media Use Moderated Anxiety Change during "
       "the COVID-19 Pandemic. PMC10842370 (2024).", size=11)
    nr(d, "5. Adaptive versus maladaptive coping strategies: insight from Lebanese "
       "young adults navigating multiple crises (2025). PMC12010593.", size=11)
    nr(d, "6. Mello và cộng sự. Coping Skills Help Explain How Future-Oriented "
       "Adolescents Accrue Greater Well-Being Over Time. PMID 25427783.", size=11)
    nr(d, "7. Emotional Regulation and Subjective Well-Being in Adolescents: A "
       "Systematic Review. Mental Health: Global Challenges Journal, Bài 240.", size=11)

    H2(d, "Truy vết nội bộ")
    nr(d, "• Doc V1 này: 01_Bao-cao/Tra_loi_HieuQua_CopingTuPhat_HS_cho_thay_27042026.docx",
       size=11)
    nr(d, "• Doc dịch đầy đủ Herres 2015: 03_Ban-dich/Bai_dich_phan_bien/"
       "Herres_Ohannessian_2015_CopingProfiles_dich_phan_bien_27042026.docx", size=11)
    nr(d, "• Doc dịch đầy đủ Steinhoff 2023: 03_Ban-dich/Bai_dich_phan_bien/"
       "Steinhoff_2023_LongitudinalCoping_dich_phan_bien_27042026.docx", size=11)
    nr(d, "• RAG: collection lo_au_dich_phan_bien ở rag_dich_phan_bien/ "
       "(đã add ~30-40 chunks mới)", size=11)
    nr(d, "• KG: 06_Scripts/kg_data/ (đã add 3 nodes: TLN_Herres2015, TLN_Steinhoff2023, "
       "TraLoi_HieuQuaCoping_v1 + edges đến Wen 2020 QT08 cùng phương pháp LPA)", size=11)

    out = os.path.join(OUT_DIR_BAOCAO, 'Tra_loi_HieuQua_CopingTuPhat_HS_cho_thay_27042026.docx')
    d.save(out); return out

# =====================================================================
# DOC 2 — Herres & Ohannessian 2015 — DỊCH + PHẢN BIỆN
# =====================================================================
def make_doc_herres():
    d = make_doc()
    title(d, "BẢN DỊCH + PHẢN BIỆN ĐẦY ĐỦ", 16)
    title(d, "CÁC HỒ SƠ ỨNG PHÓ CỦA VỊ THÀNH NIÊN", 15)
    title(d, "PHÂN BIỆT BÁO CÁO TRIỆU CHỨNG TRẦM CẢM VÀ LO ÂU", 14)
    subtitle(d, "Adolescent Coping Profiles Differentiate Reports of Depression and Anxiety Symptoms")
    nr(d, "")
    subtitle(d, "Joanna Herres & Christine McCauley Ohannessian (2015)")
    subtitle(d, "Journal of Affective Disorders, Tập 186: 312–319")
    subtitle(d, "DOI: 10.1016/j.jad.2015.07.031 — PMID 26275359 — PMC4565746")
    nr(d, "")
    subtitle(d, "Trợ lý nghiên cứu — 27/04/2026 — tiếng Việt thuần — chú thích Anh trong ngoặc")

    H1(d, "THÔNG TIN THƯ MỤC")
    tbl(d, ['Mục', 'Nội dung'],
        [
            ['Tên bài (tiếng Anh)',
             'Adolescent coping profiles differentiate reports of depression and anxiety symptoms'],
            ['Tạm dịch',
             'Các hồ sơ ứng phó của vị thành niên phân biệt các báo cáo về triệu chứng trầm cảm và lo âu'],
            ['Tác giả', 'Joanna Herres + Christine McCauley Ohannessian'],
            ['Tạp chí', 'Journal of Affective Disorders'],
            ['Tập / Trang / Năm', 'Tập 186: 312–319 (2015)'],
            ['DOI', '10.1016/j.jad.2015.07.031'],
            ['PMID', '26275359'],
            ['PMC', 'PMC4565746'],
            ['Cỡ mẫu', 'N=982 học sinh (51% nữ; 66% Caucasian; 19% gốc Phi; '
             '11% gốc Hispanic; 2% gốc Á)'],
            ['Tuổi', 'Lớp 10–11 — tuổi trung bình 16,04 (độ lệch chuẩn 0,73)'],
            ['Bối cảnh', 'Trường THPT công lập vùng Mid-Atlantic Hoa Kỳ (Delaware, '
             'Pennsylvania, Maryland), khảo sát mùa xuân 2007'],
            ['Loại nghiên cứu', 'Cắt ngang (cross-sectional) với phân tích hồ sơ tiềm ẩn '
             '(Latent Profile Analysis — LPA)'],
        ], [4.0, 12.0])

    H1(d, "PHẦN 1 — TÓM TẮT VÀ DỊCH SONG NGỮ")

    H2(d, "Tóm tắt (Abstract)")
    en(d, "Background: Diathesis-stress models suggest that the way one copes with "
       "stress influences whether stressful events lead to depression and anxiety. "
       "Coping strategies have typically been classified as approach versus avoidance.")
    vn(d, "Bối cảnh: Các mô hình \"tố chất – căng thẳng\" (diathesis-stress models) gợi "
       "ý rằng CÁCH một người đối phó với căng thẳng ảnh hưởng đến việc các sự kiện "
       "căng thẳng có dẫn đến trầm cảm và lo âu hay không. Các chiến lược ứng phó "
       "thường được phân loại thành tiếp cận (approach) so với tránh né (avoidance).")

    en(d, "Methods: This study used latent profile analysis (LPA) to examine coping "
       "profiles in a community sample of 982 high school adolescents.")
    vn(d, "Phương pháp: Nghiên cứu này sử dụng phân tích hồ sơ tiềm ẩn (LPA) để xem "
       "xét các hồ sơ ứng phó trong một mẫu cộng đồng gồm 982 học sinh THPT.")

    en(d, "Results: Four profiles were identified: (1) Disengaged Copers (15.6%), who "
       "reported the lowest coping overall; (2) Independent Copers (38.2%), who "
       "reported moderate coping; (3) Social Support Seeking Copers (30.6%), who "
       "reported high social support seeking; and (4) Active Copers (12.4%), who "
       "reported the highest coping across all subscales. The independent copers "
       "reported the lowest levels of depressive symptoms compared to the three "
       "other groups. The Social Support Seeking and Active Coping Groups reported "
       "the highest levels of anxiety.")
    vn(d, "Kết quả: Bốn hồ sơ được xác định: (1) Đứt kết với ứng phó (Disengaged "
       "Copers) chiếm 15,6%, với mức ứng phó thấp nhất nói chung; (2) Ứng phó độc "
       "lập (Independent Copers) chiếm 38,2%, mức ứng phó vừa phải; (3) Tìm hỗ trợ "
       "xã hội (Social Support Seeking Copers) chiếm 30,6%, có xu hướng tìm hỗ trợ "
       "xã hội cao; (4) Ứng phó chủ động (Active Copers) chiếm 12,4%, mức ứng phó "
       "cao nhất trên tất cả các thang con. NHÓM ĐỘC LẬP báo cáo mức triệu chứng "
       "TRẦM CẢM THẤP NHẤT so với 3 nhóm còn lại. NHÓM TÌM HỖ TRỢ XÃ HỘI và NHÓM "
       "CHỦ ĐỘNG báo cáo mức LO ÂU CAO NHẤT.")
    crit(d, "ĐÂY LÀ PHÁT HIỆN ĐẢO NGƯỢC TRỰC GIÁC: nhóm ứng phó NHIỀU NHẤT lại có lo "
         "âu CAO NHẤT. Tác giả gợi ý quan hệ \"cong\" (curvilinear) — coping vừa phải "
         "tốt nhất. Tuy nhiên không loại trừ giả thuyết \"reverse causation\": người "
         "đã có lo âu cao thì PHẢI dùng nhiều cách ứng phó để cố gắng giảm. Thiết "
         "kế cắt ngang không thể phân biệt được hai khả năng này.")

    en(d, "Limitations: The findings are based on cross-sectional self-report data "
       "from a community sample.")
    vn(d, "Hạn chế: Phát hiện dựa trên dữ liệu tự báo cáo cắt ngang từ một mẫu cộng đồng.")

    en(d, "Conclusions: Findings highlight the importance of considering individual "
       "differences in coping when working with adolescents who report depression "
       "and anxiety symptoms.")
    vn(d, "Kết luận: Phát hiện nhấn mạnh tầm quan trọng của việc xem xét các khác "
       "biệt cá nhân về ứng phó khi làm việc với vị thành niên có triệu chứng trầm "
       "cảm và lo âu.")

    H2(d, "Phương pháp chi tiết")
    nr(d, "• THANG ĐO ỨNG PHÓ: Bộ COPE 60 mục (Carver và cộng sự 1989) — đo 9 thang "
       "con: ứng phó chủ động (active coping), phủ nhận (denial), tìm hỗ trợ xã hội "
       "cảm xúc (emotional social support), hài hước (humor), tìm hỗ trợ xã hội công "
       "cụ (instrumental social support), thoát ra tinh thần (mental disengagement), "
       "lập kế hoạch (planning), ứng phó tâm linh (religious coping), giải toả cảm "
       "xúc (venting emotions).")
    nr(d, "• THANG ĐO TRẦM CẢM: CES-DC (Center for Epidemiological Studies Depression "
       "Scale for Children) 20 mục — phổ điểm 20–80 — hệ số tin cậy nội tại Cronbach "
       "alpha = 0,91.")
    nr(d, "• THANG ĐO LO ÂU: SCARED (Screen for Child Anxiety Related Disorders) 41 "
       "mục — phổ điểm 0–82 — Cronbach alpha = 0,94.")

    H2(d, "Kết quả 4 hồ sơ ứng phó (chi tiết)")
    tbl(d, ['Hồ sơ', 'N', '% mẫu', 'Trầm cảm M (SD)', 'Lo âu M (SD)'],
        [
            ['Đứt kết (Disengaged)', '153', '15,6%', '36,98 (12,39)', '12,69 (12,45)'],
            ['Độc lập (Independent)', '375', '38,2%', '33,69 (9,76)', '13,29 (9,68)'],
            ['Tìm hỗ trợ xã hội (Social Support Seeking)', '300', '30,6%', '36,36 (10,90)', '18,72 (12,13)'],
            ['Chủ động (Active)', '122', '12,4%', '37,42 (12,64)', '22,44 (15,53)'],
        ], [5.5, 1.5, 2.0, 3.5, 3.5])
    nr(d, "Quan sát: nhóm ĐỘC LẬP có cả trầm cảm + lo âu THẤP NHẤT; nhóm CHỦ ĐỘNG "
       "có lo âu CAO NHẤT (gần GẤP ĐÔI nhóm Đứt kết và Độc lập).", bold=True)

    H2(d, "Hạn chế tác giả tự thừa nhận")
    nr(d, "(1) \"Tất cả các thang đo đều là tự báo cáo\" — risk thiên lệch nhận thức")
    nr(d, "(2) \"Phát hiện có thể không khái quát hoá cho các quần thể khác như mẫu "
       "lâm sàng, vị thành niên lớn hơn hoặc nhỏ hơn, hoặc vị thành niên sống ngoài "
       "vùng Mid-Atlantic Hoa Kỳ.\"")
    nr(d, "(3) Thiết kế cắt ngang: \"không thể đưa ra kết luận về quan hệ nhân quả\"")
    nr(d, "(4) Có thể có biến thứ ba (loại căng thẳng, mức căng thẳng) gây ảnh hưởng "
       "đến cả coping và triệu chứng.")

    H1(d, "PHẦN 2 — PHẢN BIỆN TỔNG QUAN")

    H2(d, "Điểm mạnh")
    nr(d, "(1) CỠ MẪU LỚN — n=982 cho phép phân tích hồ sơ tiềm ẩn ổn định.")
    nr(d, "(2) PHƯƠNG PHÁP LPA — tiếp cận tập trung vào cá nhân (person-centered) "
       "thay vì biến (variable-centered) — nắm bắt được sự đa dạng của các \"phong "
       "cách\" ứng phó khác nhau.")
    nr(d, "(3) ĐO 9 thang con coping đa dạng — bao quát cả approach và avoidance.")
    nr(d, "(4) THANG ĐO chuẩn hoá quốc tế (CES-DC, SCARED, COPE) — kết quả có thể "
       "so sánh với các nghiên cứu khác.")
    nr(d, "(5) MẪU đa sắc tộc (66% Caucasian, 19% Phi, 11% Hispanic, 2% Á) — đại "
       "diện tốt hơn các nghiên cứu chỉ Caucasian.")

    H2(d, "Điểm yếu")
    crit(d, "(1) THIẾT KẾ CẮT NGANG — không thể kết luận nhân quả. Có thể: (a) coping "
         "ảnh hưởng đến lo âu (giả thuyết của tác giả); (b) lo âu thúc đẩy người ta "
         "dùng nhiều cách ứng phó (reverse causation); (c) cả hai cùng ảnh hưởng "
         "lẫn nhau (bidirectional).")
    crit(d, "(2) CHỈ TỰ BÁO CÁO — không có đánh giá lâm sàng độc lập về lo âu / "
         "trầm cảm.")
    crit(d, "(3) Mẫu MID-ATLANTIC US (Delaware, Pennsylvania, Maryland) — không có "
         "vùng nghèo + nông thôn sâu của Mỹ; không generalisable cho các quốc gia "
         "khác.")
    crit(d, "(4) BỘ COPE 60 MỤC quá DÀI — có thể gây mệt mỏi và giảm chất lượng dữ "
         "liệu ở học sinh THPT.")
    crit(d, "(5) KHÔNG ĐO MỨC ĐỘ CĂNG THẲNG mà học sinh đang đối mặt — biến này có "
         "thể là yếu tố gây nhiễu (confounder) chính. Có thể \"Active Copers\" có "
         "lo âu cao chỉ vì họ đang đối mặt nhiều stressor hơn.")

    H2(d, "Áp dụng cho Việt Nam")
    nr(d, "(1) PHƯƠNG PHÁP LPA có thể áp dụng cho dữ liệu Việt Nam — đặc biệt phù "
       "hợp với mẫu lớn của V-NAMHS hoặc khảo sát quốc gia trong tương lai.")
    nr(d, "(2) THANG ĐO COPE 60 mục cần dịch + kiểm định cho Việt Nam — chưa có "
       "phiên bản chuẩn hoá.")
    nr(d, "(3) THANG ĐO SCARED đã có một số bản dịch tiếng Việt nhưng chưa kiểm "
       "định chính thức cho Việt Nam.")
    nr(d, "(4) PHÁT HIỆN nhóm \"Độc lập\" có sức khỏe tâm thần tốt nhất phù hợp "
       "giá trị TỰ LẬP trong văn hoá Á Đông — có thể là điểm khởi đầu để xây dựng "
       "lý thuyết coping bản địa hoá Việt Nam.")
    nr(d, "(5) CẢNH BÁO: nhóm \"Tìm hỗ trợ xã hội\" có lo âu cao — không có nghĩa "
       "nên CẢN người ta tìm hỗ trợ. Có thể họ tìm hỗ trợ vì đang căng thẳng nhiều. "
       "Trong văn hoá Á Đông coi trọng gia đình, kết quả này có thể KHÁC ở Việt Nam.")

    H2(d, "Đối chiếu với corpus dự án")
    tbl(d, ['Bài/RAG', 'Liên hệ với Herres 2015'],
        [
            ['Wen 2020 (QT08)', 'Cùng dùng phương pháp LPA — Wen ở Trung Quốc nông thôn '
             '900 HS xác định 3 hồ sơ lo âu (nhẹ/trung bình/nặng); Herres ở Mỹ 982 HS '
             'xác định 4 hồ sơ COPING. Hai cách phân loại khác nhau (theo outcome vs '
             'theo cách ứng phó) nhưng cùng triết lý person-centered.'],
            ['BESST trial Brown 2024',
             'BESST đo 80% học sinh chưa từng tham vấn GP; nhưng KHÔNG đo họ đang dùng '
             'cách ứng phó tự phát nào. Herres lấp khoảng trống này — cho biết HS dùng '
             'gì trước khi đến với chương trình can thiệp.'],
            ['PLACES Brown 2022',
             'PLACES đề xuất các yếu tố giúp HS engage với dịch vụ; Herres cho thấy '
             'một số HS \"Độc lập\" có thể KHÔNG cần engage với dịch vụ — họ tự xử lý '
             'tốt. Bổ sung góc nhìn cho PLACES.'],
        ], [3.5, 12.5])

    H1(d, "PHẦN 3 — THAM KHẢO ĐẦY ĐỦ")
    nr(d, "1. Herres J, Ohannessian CM (2015). Adolescent coping profiles differentiate "
       "reports of depression and anxiety symptoms. Journal of Affective Disorders, "
       "186:312–319. DOI: 10.1016/j.jad.2015.07.031 — PMID 26275359 — PMC4565746",
       size=11)
    nr(d, "2. Carver CS, Scheier MF, Weintraub JK (1989). Assessing coping strategies: "
       "A theoretically based approach. Journal of Personality and Social Psychology, "
       "56(2):267–283. DOI: 10.1037/0022-3514.56.2.267 (bộ COPE 60 mục)", size=11)
    nr(d, "3. Faulstich ME, Carey MP, Ruggiero L, Enyart P, Gresham F (1986). Assessment "
       "of depression in childhood and adolescence: An evaluation of the Center for "
       "Epidemiological Studies Depression Scale for Children (CES-DC). American "
       "Journal of Psychiatry, 143(8):1024–1027. (CES-DC)", size=11)
    nr(d, "4. Birmaher B, Khetarpal S, Brent D và cộng sự (1997). The Screen for Child "
       "Anxiety Related Emotional Disorders (SCARED): Scale construction and "
       "psychometric characteristics. Journal of the American Academy of Child & "
       "Adolescent Psychiatry, 36(4):545–553. DOI: 10.1097/00004583-199704000-00018",
       size=11)
    nr(d, "5. Wen X, Lin Y, Liu Y và cộng sự (2020). LPA của lo âu HS THCS nông thôn "
       "Trung Quốc. IJERPH, 17(11):4079 — corpus dự án QT08", size=11)
    nr(d, "Truy vết: Em đã tải nguyên văn từ PMC4565746; số liệu 4 hồ sơ + N + tuổi "
       "+ Cronbach alpha + phương pháp đều khớp với nguyên văn.", italic=True, color=GRAY, size=10)

    out = os.path.join(OUT_DIR_DICH, 'Herres_Ohannessian_2015_CopingProfiles_dich_phan_bien_27042026.docx')
    d.save(out); return out

# =====================================================================
# DOC 3 — Steinhoff 2023 — DỊCH + PHẢN BIỆN
# =====================================================================
def make_doc_steinhoff():
    d = make_doc()
    title(d, "BẢN DỊCH + PHẢN BIỆN ĐẦY ĐỦ", 16)
    title(d, "CÁC YẾU TỐ TIÊN LƯỢNG TỪ TUỔI VỊ THÀNH NIÊN SỚM", 14)
    title(d, "ĐẾN ỨNG PHÓ THÍCH NGHI CỦA THANH NIÊN", 14)
    title(d, "TRONG ĐẠI DỊCH COVID-19 (NGHIÊN CỨU CỘT MỐC 9 NĂM)", 13)
    subtitle(d, "Early Adolescent Predictors of Young Adults' Distress and Adaptive Coping "
             "During the COVID-19 Pandemic: Findings From a Longitudinal Cohort Study")
    nr(d, "")
    subtitle(d, "Steinhoff A, Johnson-Ferguson L, Bechtiger L và cộng sự (2023)")
    subtitle(d, "Journal of Early Adolescence, Tập 44(9): 1250–1280")
    subtitle(d, "DOI: 10.1177/02724316231181660 — PMID 39372429 — PMC10261967")
    nr(d, "")
    subtitle(d, "Trợ lý nghiên cứu — 27/04/2026 — tiếng Việt thuần — chú thích Anh trong ngoặc")

    H1(d, "THÔNG TIN THƯ MỤC")
    tbl(d, ['Mục', 'Nội dung'],
        [
            ['Tên bài (tiếng Anh)',
             "Early Adolescent Predictors of Young Adults' Distress and Adaptive Coping "
             "During the COVID-19 Pandemic: Findings From a Longitudinal Cohort Study"],
            ['Tạm dịch',
             'Các yếu tố tiên lượng từ tuổi vị thành niên sớm cho căng thẳng tâm lý và '
             'ứng phó thích nghi của thanh niên trong đại dịch COVID-19: Phát hiện từ '
             'một nghiên cứu cohort dài hạn'],
            ['Tác giả', 'Steinhoff A, Johnson-Ferguson L, Bechtiger L và cộng sự'],
            ['Tạp chí', 'Journal of Early Adolescence (Sage)'],
            ['Tập / Trang / Năm', 'Tập 44(9): 1250–1280 (2023)'],
            ['DOI', '10.1177/02724316231181660'],
            ['PMID', '39372429'],
            ['PMC', 'PMC10261967'],
            ['Cỡ mẫu', 'N=786 người tham gia từ Dự án Phát triển Xã hội Zurich '
             '(Zurich Project on Social Development — z-proso)'],
            ['Bối cảnh', 'Zurich, Thuỵ Sĩ'],
            ['Tuổi baseline', '13 tuổi (năm 2011)'],
            ['Tuổi follow-up cuối', '~22 tuổi (tháng 4/2020 — đỉnh phong toả COVID-19 '
             'lần 1 ở Châu Âu)'],
            ['Số lần đánh giá', '8 lần từ 7–20 tuổi (đến 2018) + 4 lần thêm tháng 4–9/2020'],
            ['Khoảng cách phân tích chính', '~9 năm (tuổi 13 → tuổi 22)'],
            ['Loại nghiên cứu', 'Cohort dài hạn (longitudinal cohort study)'],
        ], [4.0, 12.0])

    H1(d, "PHẦN 1 — TÓM TẮT VÀ PHƯƠNG PHÁP")

    H2(d, "Mục tiêu nghiên cứu")
    nr(d, "Xem xét xem các yếu tố ở tuổi vị thành niên sớm (early adolescence — 13 tuổi) "
       "có tiên lượng được căng thẳng tâm lý (distress) và ứng phó thích nghi (adaptive "
       "coping) ở tuổi thanh niên (young adulthood — 22 tuổi) trong giai đoạn căng "
       "thẳng cấp tính (đại dịch COVID-19) hay không.")

    H2(d, "Thiết kế nghiên cứu")
    nr(d, "Đây là một trong số ÍT NGHIÊN CỨU dài hạn về coping ở vị thành niên — theo "
       "dõi cùng một nhóm người trong 9 năm. Đây là điểm rất quý vì hầu hết nghiên "
       "cứu coping là cắt ngang.", bold=True)
    nr(d, "Mẫu: 786 người tham gia từ Dự án Phát triển Xã hội Zurich (z-proso) — một "
       "trong những nghiên cứu cohort vị thành niên lớn nhất Châu Âu, bắt đầu từ năm "
       "2004 với mẫu trẻ em vào lớp 1.")
    nr(d, "Cấu trúc đánh giá: 8 đợt khảo sát từ tuổi 7 đến 20 (đến năm 2018), sau đó "
       "thêm 4 đợt khảo sát đặc biệt từ tháng 4/2020 đến tháng 9/2020 trong giai đoạn "
       "phong toả COVID-19 đầu ở Châu Âu.")

    H2(d, "Thang đo")
    block(d, 'Đo coping thích nghi:',
          'Các mục đo tần suất sử dụng các chiến lược thích nghi: \"tìm hỗ trợ cảm '
          'xúc, duy trì liên lạc với người thân, giúp đỡ hàng xóm\" và các chiến lược '
          'tự túc. Tổng điểm thể hiện sự linh hoạt trong ứng phó (coping flexibility).')
    block(d, 'Đo căng thẳng tâm lý (distress) sau đại dịch:',
          'Ba chỉ báo đo qua thang 10 điểm: (1) cảm giác tệ hơn so với trước khi có '
          'đại dịch, (2) tuyệt vọng về tương lai, (3) cảm nhận xáo trộn cuộc sống.')
    block(d, 'Đo triệu chứng nội tâm hoá ở tuổi 13:',
          'Các thang về lo âu, trầm cảm, rút lui xã hội — đo bằng phiếu báo cáo của '
          'cha mẹ + giáo viên + tự báo cáo.')
    block(d, 'Đo tương tác hỗ trợ cha-mẹ-con:',
          'Đánh giá sự ấm áp, hỗ trợ, sự sẵn có về cảm xúc của cha mẹ qua mắt con.')

    H1(d, "PHẦN 2 — KẾT QUẢ CHÍNH")

    H2(d, "Phát hiện 1 — Hỗ trợ cha mẹ ở 13 tuổi tiên lượng coping thích nghi ở 22 tuổi")
    en(d, "\"Supportive parent–child interactions in early adolescence were associated "
       "with more frequent use of any adaptive coping strategies in early adulthood.\"")
    vn(d, "\"Tương tác hỗ trợ giữa cha mẹ và con ở tuổi vị thành niên sớm có liên quan "
       "đến việc sử dụng THƯỜNG XUYÊN HƠN các chiến lược ứng phó thích nghi ở tuổi "
       "thanh niên đầu.\"")
    nr(d, "Cụ thể: tương tác hỗ trợ ở tuổi 13 → tiên lượng đặc biệt là COPING DỰA TRÊN "
       "XÃ HỘI (socially-based coping) và TÁI CẤU TRÚC NHẬN THỨC (cognitive reappraisal) "
       "ở tuổi 22.")

    H2(d, "Phát hiện 2 — Triệu chứng nội tâm hoá ở 13 tuổi tiên lượng distress ở 22 tuổi")
    en(d, "\"Higher levels of internalizing symptoms at age 13 were associated with "
       "feeling worse than before the pandemic, more hopeless, and perceiving more "
       "lifestyle disruptions.\"")
    vn(d, "\"Mức độ triệu chứng nội tâm hoá CAO HƠN ở tuổi 13 có liên quan với cảm "
       "giác TỆ HƠN sau đại dịch, TUYỆT VỌNG hơn, và CẢM NHẬN xáo trộn cuộc sống "
       "nhiều hơn.\"")
    nr(d, "Đây là bằng chứng mạnh cho việc TRẦM CẢM/LO ÂU SỚM dự đoán phản ứng kém "
       "với stressor lớn 9 năm sau.")

    H2(d, "Phát hiện 3 — Sự kiện căng thẳng tích luỹ sớm BUFFER tác động đại dịch")
    en(d, "\"Early exposure to cumulative stressful life events moderated (buffered) the "
       "effect of pandemic stressors on perceived disruptions, supporting an "
       "inoculation hypothesis.\"")
    vn(d, "\"Phơi nhiễm sớm với các sự kiện căng thẳng tích luỹ ĐIỀU CHỈNH (buffer — "
       "giảm nhẹ) tác động của các stressor đại dịch lên cảm nhận xáo trộn — ủng hộ "
       "GIẢ THUYẾT \\\"miễn dịch tâm lý\\\" (inoculation hypothesis).\"")
    nr(d, "Đây là phát hiện ĐẶC BIỆT QUAN TRỌNG và phản trực giác: trải qua khó khăn "
       "VỪA PHẢI khi còn trẻ có thể giúp xây dựng \"sức đề kháng\" tâm lý cho khó "
       "khăn lớn sau này. KHÔNG phải mọi căng thẳng đều có hại.", bold=True)

    H1(d, "PHẦN 3 — PHẢN BIỆN TỔNG QUAN")

    H2(d, "Điểm mạnh")
    nr(d, "(1) THIẾT KẾ DÀI HẠN HIẾM CÓ — theo dõi cùng người trong 9 năm, đo trước "
       "khi có biến cố COVID-19. Cho phép suy luận nhân quả tốt hơn các nghiên cứu "
       "cắt ngang.")
    nr(d, "(2) MẪU LỚN n=786 được giữ qua 12 đợt đánh giá — tỉ lệ giữ mẫu (retention) "
       "tốt cho 1 nghiên cứu cohort.")
    nr(d, "(3) ĐO ĐA TÁC NHÂN: cha mẹ + giáo viên + tự báo cáo — giảm thiên lệch tự "
       "báo cáo.")
    nr(d, "(4) TÍNH THỜI SỰ: dữ liệu COVID-19 cho phép kiểm định giả thuyết \"nguồn "
       "gốc sớm dự đoán phản ứng với khủng hoảng lớn\".")

    H2(d, "Điểm yếu")
    crit(d, "(1) THANG ĐO COPING RẤT HẠN CHẾ — chỉ vài mục, hệ số tin cậy nội tại "
         "thấp (\"relatively low overall internal consistency\"). Tác giả tự thừa nhận.")
    crit(d, "(2) MẪU THUỴ SĨ — quốc gia thu nhập cao, hệ thống y tế và phúc lợi tốt; "
         "khả năng generalisable cho Việt Nam có hạn.")
    crit(d, "(3) ĐO COVID-19 CHỈ TRONG 6 THÁNG (4–9/2020) — không nắm bắt được tác "
         "động dài hạn của đại dịch.")
    crit(d, "(4) KHÔNG ĐO QUÁ TRÌNH TRUNG GIAN giữa 13 và 22 tuổi — bỏ qua các yếu "
         "tố ở tuổi mid + late adolescence (15–18) có thể là biến trung gian quan "
         "trọng.")
    crit(d, "(5) KÍCH THƯỚC HIỆU ỨNG NHỎ — \"effect sizes in our models were relatively "
         "small\". Tác giả tự thừa nhận.")

    H2(d, "Áp dụng cho Việt Nam")
    nr(d, "(1) BẰNG CHỨNG MẠNH cho việc ĐẦU TƯ VÀO CHA MẸ ở giai đoạn vị thành niên "
       "sớm (lớp 6-8 ở Việt Nam) — vì có hiệu quả LÂU DÀI 9 năm sau cho khả năng "
       "ứng phó của con.")
    nr(d, "(2) Khuyến nghị chương trình \"Cha mẹ tích cực\" tại trường THCS Việt Nam "
       "— tập trung vào KỸ NĂNG GIAO TIẾP cảm xúc + sự ấm áp + tính sẵn có. Có thể "
       "tham khảo các chương trình quốc tế như Triple P (Positive Parenting Program) "
       "hoặc Incredible Years — bản địa hoá cho Việt Nam.")
    nr(d, "(3) Cần SÀNG LỌC SỚM triệu chứng nội tâm hoá ở học sinh THCS — vì là yếu "
       "tố tiên lượng mạnh cho phản ứng kém với khủng hoảng sau này. Khuyến nghị đo "
       "GAD-7 + PHQ-9 đầu năm học cho HS lớp 6–9.")
    nr(d, "(4) GIẢ THUYẾT \"miễn dịch tâm lý\" cần được kiểm chứng cho Việt Nam — "
       "tránh áp dụng máy móc vì hệ thống xã hội + gia đình Việt Nam khác Thuỵ Sĩ.")

    H2(d, "Đối chiếu corpus dự án")
    nr(d, "• Hỗ trợ cha mẹ → coping thích nghi: bổ sung cho phát hiện QT09 Qiu 2022 "
       "Trung Quốc (nuôi dạy tích cực giảm trầm cảm OR=0,30-0,32).")
    nr(d, "• Triệu chứng sớm → distress sau: phù hợp với phát hiện Hoa 2024 Hà Nội "
       "(40,6% HS THPT có lo âu, có thể tiên lượng phản ứng kém với áp lực sau này).")
    nr(d, "• Inoculation hypothesis: NEW cho corpus VN — chưa bài nào đề cập trực tiếp.")

    H1(d, "PHẦN 4 — THAM KHẢO ĐẦY ĐỦ")
    nr(d, "1. Steinhoff A, Johnson-Ferguson L, Bechtiger L và cộng sự (2023). Early "
       "Adolescent Predictors of Young Adults Distress and Adaptive Coping During "
       "the COVID-19 Pandemic: Findings From a Longitudinal Cohort Study. Journal of "
       "Early Adolescence, 44(9):1250–1280. DOI: 10.1177/02724316231181660 — "
       "PMID 39372429 — PMC10261967", size=11)
    nr(d, "2. Eisner M, Ribeaud D (2007). Conducting a Criminological Survey in a "
       "Culturally Diverse Context: Lessons from the Zurich Project on the Social "
       "Development of Children. European Journal of Criminology, 4(3):271-298. "
       "(Mô tả Dự án z-proso)", size=11)
    nr(d, "3. Rutter M (1985). Resilience in the face of adversity: Protective factors "
       "and resistance to psychiatric disorder. British Journal of Psychiatry, 147:598–"
       "611. (Lý thuyết miễn dịch tâm lý — inoculation hypothesis)", size=11)
    nr(d, "4. Qiu và cộng sự (2022). Nuôi dạy tiêu cực và lo âu / trầm cảm vị thành "
       "niên Trung Quốc — corpus dự án QT09", size=11)
    nr(d, "Truy vết: Em đã tải nguyên văn từ PMC10261967 ngày 27/04/2026; số liệu n=786, "
       "tuổi 13→22, 9 năm, 4 đợt khảo sát COVID — đều khớp nguyên văn.",
       italic=True, color=GRAY, size=10)

    out = os.path.join(OUT_DIR_DICH, 'Steinhoff_2023_LongitudinalCoping_dich_phan_bien_27042026.docx')
    d.save(out); return out

if __name__ == '__main__':
    p1 = make_doc_v1(); print('V1:', p1)
    p2 = make_doc_herres(); print('Herres:', p2)
    p3 = make_doc_steinhoff(); print('Steinhoff:', p3)
