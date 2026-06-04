# -*- coding: utf-8 -*-
"""Dịch song ngữ + Phản biện đầy đủ bài Brown 2022 PLACES.
Bài gốc: Brown JSL, Lisk S, Carter B, Stevelink SAM, Van Lieshout R, Michelson D (2022).
How Can We Actually Change Help-Seeking Behaviour for Mental Health Problems among the
General Public? Development of the 'PLACES' Model.
IJERPH, 19(5):2831. doi:10.3390/ijerph19052831 (PMC8909998)
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\03_Ban-dich\Bai_dich_phan_bien\PLACES_Brown_2022_dich_phan_bien_25042026.docx'

RED   = RGBColor(0xC0, 0x00, 0x00)
BLUE  = RGBColor(0x00, 0x70, 0xC0)
GRAY  = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0x00, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name = 'Times New Roman'; s.font.size = Pt(13)
s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.4
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.5); sec.right_margin = Cm(2.0)

def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = d.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def title(text, size=18):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(size); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def subtitle(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def H1(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(15); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def H2(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True; r.font.size = Pt(13); r.font.color.rgb = BLUE
    r.font.name = 'Times New Roman'
def en(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.4)
    r = p.add_run(text); r.italic = True; r.font.size = Pt(11); r.font.color.rgb = GRAY
    r.font.name = 'Times New Roman'
def vn(text):
    p = d.add_paragraph()
    r = p.add_run(text); r.font.size = Pt(13); r.font.name = 'Times New Roman'
def crit(text):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('[Phản biện] '); r.bold = True; r.font.color.rgb = RED; r.font.size = Pt(12); r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.font.color.rgb = RED; r2.font.size = Pt(12); r2.font.name = 'Times New Roman'
def note(text, color=GREEN):
    p = d.add_paragraph(); p.paragraph_format.left_indent = Cm(0.5)
    r = p.add_run('[Ghi chú dịch giả] '); r.bold = True; r.italic = True; r.font.color.rgb = color; r.font.size = Pt(11); r.font.name = 'Times New Roman'
    r2 = p.add_run(text); r2.italic = True; r2.font.color.rgb = color; r2.font.size = Pt(11); r2.font.name = 'Times New Roman'
def nr(text, bold=False, size=13, color=None, italic=False):
    p = d.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.italic = italic
    if color is not None: r.font.color.rgb = color

# =====================================================================
# TRANG BÌA
# =====================================================================
title("BẢN DỊCH SONG NGỮ + PHẢN BIỆN", 16)
title("LÀM SAO ĐỂ THỰC SỰ THAY ĐỔI HÀNH VI", 16)
title("TÌM KIẾM TRỢ GIÚP SỨC KHỎE TÂM THẦN", 16)
title("MÔ HÌNH 'PLACES'", 17)
subtitle("How Can We Actually Change Help-Seeking Behaviour for Mental Health Problems")
subtitle("among the General Public? Development of the 'PLACES' Model")
nr("")
subtitle("June S. L. Brown, Stephen Lisk, Ben Carter, Sharon A. M. Stevelink,")
subtitle("Ryan Van Lieshout & Daniel Michelson (2022)")
subtitle("International Journal of Environmental Research and Public Health, 19(5): 2831")
subtitle("doi: 10.3390/ijerph19052831 — PMC8909998 — Open Access (CC BY 4.0)")
subtitle("Loại bài: ARTICLE — Mô hình lý thuyết tổng hợp + 4 case studies thực nghiệm")
nr("")
nr("Trợ lý nghiên cứu — Dịch + Phản biện — 25/04/2026", italic=True, color=GRAY, size=10)
nr("Quy trình: dịch song ngữ EN↔VN + phản biện chữ đỏ có dẫn chứng từ corpus 35+ bài + "
   "đối chiếu PMC NCBI verbatim + so sánh với mô hình lý thuyết khác (HBM, TPB, Andersen). "
   "Kiểm 3 vòng theo memory feedback_research_workflow.md.",
   italic=True, color=GRAY, size=10)

# =====================================================================
# THÔNG TIN THƯ MỤC
# =====================================================================
H1("THÔNG TIN THƯ MỤC")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài (EN)', 'How Can We Actually Change Help-Seeking Behaviour for Mental '
                          'Health Problems among the General Public? Development of the '
                          '"PLACES" Model'],
        ['Tên bài (VN)', 'Làm sao để thực sự thay đổi hành vi tìm kiếm trợ giúp SKTT trong '
                          'cộng đồng? Phát triển mô hình "PLACES"'],
        ['Tác giả',
         '1. June S. L. Brown — Psychology Department, IoPPN, KCL Denmark Hill (corresponding)\n'
         '2. Stephen Lisk — Psychology Department, IoPPN, KCL\n'
         '3. Ben Carter — Biostatistics Department, IoPPN, KCL\n'
         '4. Sharon A. M. Stevelink — Department of Psychological Medicine, IoPPN, KCL\n'
         '5. Ryan Van Lieshout — Department of Psychiatry & Behavioural Neurosciences, '
         'McMaster University, Canada\n'
         '6. Daniel Michelson — Psychology Department, Sussex University, Brighton'],
        ['Liên hệ', 'June.Brown@kcl.ac.uk'],
        ['Tạp chí', 'International Journal of Environmental Research and Public Health (IJERPH) — MDPI'],
        ['Tập / Số / Bài', 'Vol. 19, No. 5, Article 2831 (2022)'],
        ['DOI', '10.3390/ijerph19052831'],
        ['PubMed Central', 'PMC8909998'],
        ['Loại bài', 'Article (paper lý thuyết + 4 case studies)'],
        ['Ngày nhận / chấp nhận / công bố', '20/12/2021 → 31/01/2022 → 28/02/2022'],
        ['Số trang', '12 trang (PDF)'],
        ['License', 'Creative Commons CC BY 4.0 (Open Access)'],
        ['Funding', 'Phối hợp nhiều grant: PhD studentship Birmingham (Brown 2000); Guy\'s '
                    '& St Thomas\' (Brown 2004); NIHR RfPB PB-PG-1207-15154 (Horrell 2014); '
                    'NIHR RfPB PB-PG-0213-30072 (Brown 2019); Canada Research Chairs + CIHR '
                    'COVID-19 #450207 (Van Lieshout)'],
        ['Conflict of interest', 'Không có'],
        ['Mã trong corpus dự án', 'Chưa có mã chính thức (chưa được index trong RAG/KG); '
                                  'có liên kết qua glossary BESST/PLACES'],
    ], [4.5, 11.5])

H2("Định vị bài này trong tổng quan dự án")
vn("Đây là PAPER LÝ THUYẾT + 4 CASE STUDIES quan trọng do nhóm Brown (King's College "
   "London) công bố năm 2022 — cung cấp KHUNG TƯ DUY (PLACES model) làm nền tảng cho "
   "thử nghiệm BESST 2024 sau này. Khác với editorial Brown & Carter 2025 (B5), đây là "
   "PEER-REVIEWED ARTICLE với 4 trường hợp thực nghiệm (stress workshop, depression "
   "workshop, adolescent stress workshop = DISCOVER, postnatal depression workshop) — "
   "tổng hợp lại thành mô hình 6 yếu tố P-L-A-C-E-S.")
crit("Loại bằng chứng: Đây là article PEER-REVIEWED nhưng KHÔNG phải SR/MA/RCT. Cấu trúc "
     "bài là \"developmental paper\" — tổng hợp 4 case studies độc lập của chính nhóm "
     "tác giả thành một mô hình. 4 case studies đều là quasi-experimental (không có "
     "control group song song). Theo Oxford CEBM 2011: case series = level 4 evidence. "
     "Vì vậy mô hình PLACES là HEURISTIC / THEORETICAL FRAMEWORK chứ chưa phải intervention "
     "có RCT validation. Tới BESST 2024 mới có một phần PLACES (cụ thể là Self-referral) "
     "được test trong cluster RCT — và cho effect size SMALL (d=−0,17). PLACES đầy đủ "
     "(cả 6 yếu tố) chưa từng được test independently.")

# =====================================================================
# PHẦN 1 — DỊCH SONG NGỮ
# =====================================================================
H1("PHẦN 1 — DỊCH SONG NGỮ ĐẦY ĐỦ")

H2("Tóm tắt (Abstract)")
en("Good treatment uptake is essential for clinically effective interventions to be fully "
   "utilised. Numerous studies have examined barriers to help-seeking for mental health "
   "treatment and to a lesser extent, facilitators. However, much of the current research "
   "focuses on changing help-seeking attitudes, which often do not lead to changes in "
   "behaviour. There is a clear gap in the literature for interventions that successfully "
   "change help-seeking behaviour among the general public. This gap is particularly "
   "relevant for early intervention. Here we describe the development of a new model "
   "which combines facilitators to treatment and an engaging, acceptable intervention for "
   "the general public. It is called the 'PLACES' (Publicity, Lay, Acceptable, Convenient, "
   "Effective, Self-referral) model of treatment engagement. It is based on theoretical "
   "work, as well as empirical research on a low intensity psychoeducational cognitive "
   "behavioural therapy (CBT) intervention: one-day workshops for stress and depression. "
   "In this paper, we describe the development of the model and the results of its use "
   "among four different clinical groups (adults experiencing stress, adults experiencing "
   "depression, adolescents (age 16–18) experiencing stress, and mothers with postnatal "
   "depression). We recorded high rates of uptake by people who have previously not "
   "sought help and by racial and ethnic minority groups across all four of these "
   "clinical groups. The clinical and research implications and applications of this "
   "model are discussed.")
vn("Tỉ lệ tiếp nhận điều trị tốt là điều thiết yếu để các can thiệp hiệu quả lâm sàng "
   "được sử dụng đầy đủ. Nhiều nghiên cứu đã xem xét các RÀO CẢN của việc tìm kiếm trợ "
   "giúp SKTT, và ít hơn là các YẾU TỐ THÚC ĐẨY (facilitators). Tuy nhiên, phần lớn "
   "nghiên cứu hiện nay tập trung vào thay đổi THÁI ĐỘ tìm trợ giúp — vốn thường KHÔNG "
   "dẫn đến thay đổi HÀNH VI. Có một khoảng trống rõ rệt trong y văn về các can thiệp "
   "thực sự thay đổi được hành vi tìm trợ giúp ở cộng đồng. Khoảng trống này đặc biệt "
   "quan trọng đối với can thiệp sớm. Tại đây, chúng tôi mô tả việc phát triển một mô "
   "hình mới — kết hợp các facilitator của điều trị với một can thiệp thu hút, có thể "
   "chấp nhận được cho cộng đồng. Mô hình được gọi là 'PLACES' (Publicity, Lay, "
   "Acceptable, Convenient, Effective, Self-referral) của treatment engagement. Mô hình "
   "dựa trên công trình lý thuyết cũng như nghiên cứu thực nghiệm về một can thiệp CBT "
   "tâm lý giáo dục cường độ thấp: workshop 1 ngày cho stress và trầm cảm. Trong bài "
   "này, chúng tôi mô tả việc phát triển mô hình và kết quả của việc sử dụng nó trên 4 "
   "nhóm lâm sàng khác nhau (người lớn trải nghiệm stress, người lớn trầm cảm, vị thành "
   "niên 16–18 tuổi trải nghiệm stress, và mẹ trầm cảm sau sinh). Chúng tôi ghi nhận "
   "tỷ lệ tiếp nhận CAO ở những người TRƯỚC ĐÂY chưa từng tìm trợ giúp và ở các nhóm "
   "thiểu số chủng tộc / sắc tộc, ở cả 4 nhóm lâm sàng. Các hệ quả và ứng dụng lâm sàng, "
   "nghiên cứu của mô hình này được thảo luận.")
crit("Abstract trung thực ở chỗ thừa nhận \"khoảng trống\" giữa attitude và behaviour. "
     "Nhưng cũng có 2 điều cần cảnh báo: (1) \"high rates of uptake\" — không có "
     "ngưỡng so sánh / không control group → khó định lượng \"high\" là cao đến đâu so "
     "với baseline. (2) Mô hình được PHÁT TRIỂN từ 4 case studies → có nguy cơ "
     "OVERFITTING — yếu tố nào hữu ích trong 4 case này có thể không generalise sang "
     "context khác (vd. VN, LMIC). Cần validation độc lập.")

# ---------- 1. Introduction ----------
H2("1. Mở đầu (Introduction)")
en("It is estimated that 17% of the adult population in England have mental health "
   "problems. However, only 30% of those affected individuals seek professional help. "
   "Problems with low rates of help-seeking have been widely described and can result in "
   "poorer mental health outcomes, including increased chronicity. Low help-seeking also "
   "leads to poorer recruitment and selection bias in research studies, raising the "
   "possibility that these findings may lack external validity.")
vn("Ước tính 17% dân số trưởng thành ở Anh có vấn đề SKTT. Tuy nhiên, chỉ 30% trong số "
   "những người bị ảnh hưởng tìm kiếm trợ giúp chuyên nghiệp. Vấn đề tỉ lệ tìm trợ giúp "
   "thấp đã được mô tả rộng rãi và có thể dẫn đến kết cục SKTT kém hơn — bao gồm tăng "
   "tính mạn tính. Tỉ lệ tìm trợ giúp thấp cũng dẫn đến tuyển mẫu kém và sai lệch chọn "
   "mẫu trong các nghiên cứu — làm tăng khả năng các phát hiện thiếu tính giá trị bên "
   "ngoài (external validity).")
crit("Số 17% (UK adult mental health) khớp với khảo sát APMS 2014 (McManus). Số 30% "
     "tìm trợ giúp khá cao so với VN (8,4%, V-NAMHS). Đây cũng là rationale rất mạnh "
     "cho VN: nếu UK với hệ thống NHS phát triển mà chỉ 30% tìm trợ giúp, thì VN với "
     "hệ thống y tế công yếu hơn cần INTERVENTION TĂNG HELP-SEEKING khẩn cấp.")

en("Typically, research focuses on barriers and facilitators regarding attitudes around "
   "help-seeking. However, studies have had mixed success regarding interventions to "
   "improve help-seeking behaviour. A systematic review that was conducted assessed the "
   "success of randomised controlled trials (RCT) of interventions aimed at increasing "
   "help-seeking behaviour in adults with problems of depression, anxiety, and general "
   "distress. Mental health literacy was the only intervention found to be associated "
   "with improved help-seeking attitudes, but this had no effect on help-seeking "
   "behaviour. The authors concluded, \"signifi... very little is known about what "
   "interventions increase help-seeking behaviour\".")
vn("Thông thường, nghiên cứu tập trung vào RÀO CẢN và YẾU TỐ THÚC ĐẨY liên quan đến THÁI "
   "ĐỘ về help-seeking. Tuy nhiên, các nghiên cứu cho kết quả PHA TRỘN về can thiệp "
   "thực sự cải thiện hành vi help-seeking. Một tổng quan hệ thống đã đánh giá thành "
   "công của các RCT can thiệp nhằm tăng hành vi help-seeking ở người lớn có vấn đề "
   "trầm cảm, lo âu, và đau khổ chung. Mental health literacy là can thiệp DUY NHẤT "
   "được phát hiện có liên quan đến cải thiện THÁI ĐỘ help-seeking — nhưng KHÔNG có "
   "tác dụng lên HÀNH VI help-seeking. Các tác giả kết luận: \"đáng chú ý là... rất ít "
   "được biết về can thiệp nào thực sự tăng hành vi help-seeking\".")
crit("Đây là lập luận then chốt cho sự tồn tại của PLACES — gap giữa thái độ và hành "
     "vi (attitude-behaviour gap, kinh điển trong tâm lý học xã hội). Mental Health "
     "Literacy chỉ chạm tới attitude. Đây là điểm quan trọng — VN cũng cần cảnh báo "
     "khi triển khai literacy chương trình: đo cả hành vi (vd. số HS đăng ký dịch vụ "
     "tư vấn học đường sau intervention) chứ không chỉ kiến thức/thái độ.")

en("Another systematic review examined the effectiveness of help-seeking interventions "
   "in changing attitudes, intentions, and behaviour. Differences in help-seeking "
   "behaviour were found, but only among those with mental illness and those at risk of "
   "mental illness. However, no changes in help-seeking behaviour were found among the "
   "general public or among children or adolescents. This is important because they will "
   "often experience some troubling symptoms but may be unsure about what these mean, "
   "may try to cope with the problems themselves or may feel reluctant to seek help "
   "from a professional.")
vn("Một tổng quan hệ thống khác đã xem xét hiệu quả của các can thiệp help-seeking trong "
   "thay đổi thái độ, ý định, và hành vi. Có thấy khác biệt về hành vi help-seeking, "
   "nhưng CHỈ ở những người ĐÃ CÓ bệnh tâm thần và những người CÓ NGUY CƠ. Tuy nhiên, "
   "KHÔNG thấy thay đổi hành vi help-seeking ở cộng đồng chung, hoặc ở trẻ em/vị thành "
   "niên. Điều này quan trọng vì họ thường trải qua một số triệu chứng đáng lo ngại "
   "nhưng có thể không chắc chắn về ý nghĩa của chúng, có thể tự đối phó, hoặc có thể "
   "ngại tìm trợ giúp từ chuyên gia.")
crit("Đây là điểm phản biện quan trọng cho VN: nếu can thiệp help-seeking ở UK còn "
     "không hiệu quả với cộng đồng chung và VTN, thì VN với context khác càng phải có "
     "RCT riêng — không thể giả định mô hình UK transferable. Đặc biệt vì văn hoá Á "
     "Đông có \"face culture\" — kỳ thị nặng hơn UK.")

en("This demonstrates a major gap in the literature regarding interventions aiming to "
   "improve help-seeking behaviour among the general public, specifically adults and "
   "children or adolescents. It is known that the onset of mental health problems for "
   "75% of people occurs by age 24. To prevent these problems from arising or becoming "
   "chronic, it is vital we offer early intervention when symptoms are first experienced, "
   "whilst at school, college, or among people who may be unsure of their mental health "
   "status.")
vn("Điều này thể hiện một KHOẢNG TRỐNG LỚN trong y văn về các can thiệp nhằm cải thiện "
   "hành vi help-seeking ở cộng đồng — đặc biệt là người lớn và trẻ em/VTN. Ta biết "
   "rằng khởi phát các vấn đề SKTT ở 75% người là trước tuổi 24. Để ngăn các vấn đề "
   "này xuất hiện hoặc trở thành mạn tính, việc cung cấp CAN THIỆP SỚM khi triệu chứng "
   "lần đầu xuất hiện — tại trường học, đại học, hoặc cho những người không chắc chắn "
   "về tình trạng SKTT của mình — là tối quan trọng.")

en("Therefore, in this paper, we present a model we have been using which describes how "
   "members of the public with problems can more easily access professional services. In "
   "developing this model, we were informed by the conceptual model of help-seeking by "
   "Gask, which uses a patient-centred approach. The \"Community engagement\" aspect of "
   "this model is particularly important as it encompasses the uncertainty of individuals "
   "first engaging in seeking help with processes such as 'candidacy' (should I be "
   "seeking help?), 'navigation' (where do I seek help from?) and 'appearance' (actually "
   "going to seek help). The delivery of tailored psychosocial interventions is equally "
   "important as these need to fit with the needs of a population (e.g., ethnic "
   "minorities) and the social and cultural aspects of the group.")
vn("Vì vậy, trong bài này chúng tôi trình bày một mô hình đã được sử dụng — mô tả cách "
   "thành viên cộng đồng có vấn đề có thể tiếp cận dễ dàng hơn các dịch vụ chuyên môn. "
   "Trong việc phát triển mô hình này, chúng tôi đã tham khảo mô hình khái niệm về "
   "help-seeking của Gask — sử dụng cách tiếp cận lấy bệnh nhân làm trung tâm. Khía cạnh "
   "\"Community engagement\" của mô hình này đặc biệt quan trọng — bao quát sự không "
   "chắc chắn của cá nhân khi lần đầu engage với việc tìm trợ giúp qua các quá trình "
   "như: 'candidacy' (mình có nên tìm trợ giúp không?), 'navigation' (mình tìm trợ "
   "giúp ở đâu?), và 'appearance' (thực sự đi tìm trợ giúp). Việc cung cấp các can thiệp "
   "tâm lý xã hội được điều chỉnh (tailored) cũng quan trọng tương đương — cần phù hợp "
   "với nhu cầu của một nhóm dân số (vd. nhóm thiểu số chủng tộc) và các khía cạnh xã "
   "hội, văn hoá của nhóm.")
crit("Mô hình của Gask 2012 (Behavioural Model of Health Services use BMHS) là cơ sở lý "
     "thuyết quan trọng của PLACES. Cùng họ này còn có: Andersen Behavioral Model 1995 "
     "(còn được gọi là Andersen-Newman 1973), Health Belief Model (Rosenstock 1974), "
     "Theory of Planned Behaviour (Ajzen 1991). PLACES vay mượn từ Gask nhưng đơn giản "
     "hơn — chỉ tập trung vào treatment engagement (không cover policy hoặc "
     "infrastructure). Điều này vừa là điểm mạnh (operational) vừa là điểm yếu (bỏ "
     "qua structural barriers).")

en("Help-seeking is largely researched in relation to barriers rather than facilitators, "
   "with stigma being the most commonly researched barrier to help-seeking. Despite its "
   "attention, it is actually listed as only the 4th most common barrier. However, the "
   "authors note that stigma is also highly likely to influence other barriers.")
vn("Help-seeking phần lớn được nghiên cứu liên quan đến rào cản hơn là yếu tố thúc đẩy, "
   "với KỲ THỊ là rào cản được nghiên cứu nhiều nhất. Dù được chú ý nhiều, kỳ thị thực "
   "ra chỉ được xếp là rào cản phổ biến THỨ 4. Tuy nhiên, các tác giả lưu ý kỳ thị "
   "cũng có khả năng cao ảnh hưởng đến các rào cản khác.")
crit("Đây là PHÁT HIỆN BẤT NGỜ và quan trọng: kỳ thị tuy được nói nhiều nhất NHƯNG xếp "
     "thứ 4 về tần suất rào cản. Theo Clement 2014 (Psychol Med — đã trích trong PLACES "
     "ref 9): các rào cản phổ biến hơn kỳ thị bao gồm \"perceived need\" (tự nhận xét "
     "có cần trợ giúp không), \"self-reliance\" (tự xử), và \"navigation\" (không biết "
     "tìm ở đâu). Áp dụng VN: chỉ tập trung giảm stigma không đủ — cần đồng thời tăng "
     "perceived need (qua literacy) và navigation (qua self-referral system rõ ràng).")

en("How individuals view their mental health problems is key, and low perceived need is "
   "another very common barrier. Recent work conducted in a military population found "
   "that the perceived need for treatment (and not stigma) was the most prominent barrier "
   "to care. Self-reliance is often found to be the preferred way of coping, particularly "
   "among young people.")
vn("Cách cá nhân NHÌN NHẬN vấn đề SKTT của mình là then chốt, và việc \"thấy không cần "
   "thiết\" (low perceived need) là một rào cản rất phổ biến khác. Công trình gần đây "
   "trên một mẫu quân nhân cho thấy perceived need (chứ không phải stigma) là rào cản "
   "NỔI BẬT NHẤT đối với chăm sóc. Tự dựa vào chính mình (self-reliance) thường được "
   "phát hiện là cách đối phó được ưa chuộng, đặc biệt ở người trẻ.")

en("Poor mental health literacy invariably affects help-seeking; this refers to a limited "
   "understanding of mental health problems as well as not knowing where to go for help. "
   "A related problem is how the general public first seeks help: going to their general "
   "practitioner (GP) has been the standard way of accessing help in many countries but "
   "may act as a barrier for some racial and ethnic groups.")
vn("Mental health literacy kém không tránh khỏi ảnh hưởng tới help-seeking — bao gồm hiểu "
   "biết hạn chế về vấn đề SKTT cũng như không biết tìm trợ giúp ở đâu. Một vấn đề liên "
   "quan là cách cộng đồng đầu tiên tìm trợ giúp: đến bác sĩ đa khoa (GP) là cách chuẩn "
   "ở nhiều quốc gia, nhưng có thể là rào cản cho một số nhóm chủng tộc/sắc tộc.")
crit("Áp dụng VN: \"GP\" tương đương trạm y tế xã/phường — VTN VN gần như KHÔNG bao giờ "
     "đến trạm y tế cho vấn đề SKTT. Cũng không có hệ thống GP gia đình. Vì vậy "
     "self-referral qua trường học là đặc biệt quan trọng cho VN — không thể giả định "
     "kênh GP UK hoạt động.")

en("Finally, individuals may view treatment as not being very acceptable and may not take "
   "it up and/or drop out as a result. For example, online interventions (especially if "
   "self-guided) have this problem. Services can also be inconvenient when weekly "
   "sessions are offered during office hours between 9–5 pm and/or if delivered in "
   "formalised mental health settings.")
vn("Cuối cùng, cá nhân có thể nhìn nhận điều trị là KHÔNG ĐỦ CHẤP NHẬN ĐƯỢC và do đó "
   "không tiếp nhận và/hoặc bỏ giữa chừng. Ví dụ: can thiệp online (đặc biệt khi tự "
   "hướng dẫn) có vấn đề này. Dịch vụ cũng có thể KHÔNG THUẬN TIỆN khi các buổi hàng "
   "tuần được cung cấp trong giờ hành chính 9-17h và/hoặc tại các cơ sở SKTT chính thức.")

en("The 'PLACES' model was initially developed in the context of developing large-scale "
   "stress workshops for a city-wide mental health promotion campaign and then extended "
   "with depression workshops. It was further tested with adolescents experiencing stress "
   "and mothers with postnatal depression. The interventions have also been shown to be "
   "effective and are separately reported for stress, depression, adolescents, and with "
   "mothers affected by postnatal depression.")
vn("Mô hình 'PLACES' ban đầu được phát triển trong bối cảnh xây dựng các workshop stress "
   "quy mô lớn cho một chiến dịch quảng bá SKTT toàn thành phố — và sau đó mở rộng với "
   "các workshop trầm cảm. Sau đó được kiểm thử thêm với VTN trải qua stress và các bà "
   "mẹ trầm cảm sau sinh. Các can thiệp này cũng đã được chứng minh hiệu quả và được "
   "báo cáo riêng biệt cho stress, trầm cảm, VTN, và mẹ bị trầm cảm sau sinh.")

# ---------- 2. The Development and Rationale ----------
H2("2. Phát triển và Cơ sở của các Yếu tố Thúc đẩy trong Mô hình PLACES")

en("This is a summary paper of the treatment engaging factors that we have developed "
   "and utilised. We aim to synthesise our studies covering the different clinical "
   "areas, as well as different sociodemographic groups, into one paper. This paper will "
   "report on how successful the model has been at (a) engaging total numbers of "
   "participants, (b) attracting non-consulters who had not consulted their general "
   "practitioners (GPs) or professional services, (c) those with severe problems, and "
   "(d) engaging those from racial and ethnic minority groups.")
vn("Đây là một paper tổng hợp các yếu tố thu hút điều trị mà chúng tôi đã phát triển "
   "và sử dụng. Chúng tôi nhằm tổng hợp các nghiên cứu của mình bao gồm các lĩnh vực "
   "lâm sàng và các nhóm xã hội-nhân khẩu khác nhau thành một paper. Bài này sẽ báo "
   "cáo về sự thành công của mô hình ở: (a) thu hút TỔNG SỐ người tham gia, (b) thu "
   "hút \"non-consulters\" — những người chưa từng tham vấn GP hay dịch vụ chuyên "
   "nghiệp, (c) những người có vấn đề NẶNG, và (d) thu hút những người từ nhóm thiểu "
   "số chủng tộc/sắc tộc.")
crit("Đây là 4 chỉ số đánh giá thành công của PLACES theo các tác giả. Chú ý: KHÔNG có "
     "chỉ số nào đo CLINICAL OUTCOME (giảm triệu chứng) — chỉ đo ENGAGEMENT. Đây là "
     "điểm quan trọng: PLACES được thiết kế để giải quyết \"engagement gap\" chứ không "
     "phải \"clinical effectiveness gap\". Hai vấn đề khác nhau. Khi áp dụng VN cần phân "
     "biệt rõ.")

# ---------- 2.1 Stress workshops (adults) ----------
H2("2.1. Case 1 — Workshop Stress (người lớn)")
en("The large-scale workshops were part of a city-wide mental health promotion campaign. "
   "The aims of the workshops were: (1) To be as accessible as possible, especially to "
   "non-consulters. (2) To offer an acceptable large-scale intervention.")
vn("Các workshop quy mô lớn này là một phần của chiến dịch quảng bá SKTT toàn thành "
   "phố. Mục tiêu của workshop là: (1) Càng tiếp cận được càng tốt, đặc biệt cho "
   "non-consulters. (2) Cung cấp một can thiệp quy mô lớn có thể chấp nhận được.")

en("The stress workshop intervention was a 'low-intensity', large-scale, psychoeducational "
   "day-long cognitive behavioural therapy (CBT) workshop for up to 30 people. It was "
   "designed to be brief and as clinically accessible as possible, with few barriers to "
   "help-seeking. The content of the programme was informed by the evidence-based "
   "principles of CBT. Because of the possible stigmatising effects of diagnostic labels, "
   "the wording of the programme (and its publicity) avoided the use of such words as "
   "'anxiety' and 'depression'. Instead, words such as 'stress' were used.")
vn("Can thiệp workshop stress là một workshop CBT cường độ thấp (low-intensity), quy "
   "mô lớn, có tính tâm lý giáo dục, kéo dài 1 ngày, cho tối đa 30 người. Được thiết "
   "kế để NGẮN và càng dễ tiếp cận lâm sàng càng tốt, với ít rào cản help-seeking. Nội "
   "dung chương trình dựa trên các nguyên lý CBT có bằng chứng. Vì các tác động kỳ thị "
   "có thể có của nhãn chẩn đoán, ngôn ngữ chương trình (và tài liệu quảng cáo) TRÁNH "
   "sử dụng các từ như 'lo âu' và 'trầm cảm'. Thay vào đó, các từ như 'stress / căng "
   "thẳng' được sử dụng.")
crit("CHIẾN LƯỢC NGÔN NGỮ này là then chốt và rất khả thi cho VN. Đề xuất tương đương: "
     "thay \"workshop về lo âu/trầm cảm\" → \"workshop quản lý căng thẳng học tập\" / "
     "\"phòng học kỹ năng đối phó áp lực\" / \"khoá học tự tin và bình tâm\". Đây là "
     "một trong những bài học transferable nhất của PLACES.")

en("Based on the literature, it was decided to use four methods to engage the public: "
   "self-referral (S), publicity (P), acceptable (and engaging) intervention (A), and "
   "convenient location (C). Figure 1 describe how these methods relate to common "
   "barriers accessing services, namely mental health literacy, stigma, and structural "
   "barriers, as well as intervention barriers such as the acceptability and convenience "
   "of the intervention.")
vn("Dựa trên y văn, các tác giả quyết định sử dụng 4 PHƯƠNG PHÁP để thu hút cộng đồng: "
   "self-referral (S), publicity (P), can thiệp acceptable (và thu hút) (A), và vị "
   "trí convenient (C). Hình 1 mô tả các phương pháp này liên hệ với các rào cản phổ "
   "biến trong tiếp cận dịch vụ — gồm mental health literacy, stigma, và structural "
   "barriers — cũng như các rào cản can thiệp như acceptability và convenience của can "
   "thiệp.")
note("Hình 1 trong PDF gốc (trang 3) là sơ đồ kết nối 6 yếu tố P-L-A-C-E-S với 3 nhóm "
     "rào cản (mental health literacy, stigma, structural barriers) và acceptability "
     "of psychosocial interventions. Đây là bản đồ khái niệm trung tâm của bài.")

H2("Yếu tố 1 — SELF-REFERRAL (S)")
en("GP referrals have been the standard way of accessing help in many countries. However, "
   "there are significant problems with access regarding some racial and ethnic minority "
   "groups (e.g., South Asian, black Caribbean, African) who are often reluctant to "
   "consult their GP. A self-referral route was therefore developed for the workshops "
   "to allow easier access. This involved publicising the workshops widely in different "
   "community settings, with contact details highlighted to enable direct access.")
vn("Giới thiệu qua GP là cách chuẩn ở nhiều quốc gia. Tuy nhiên có những vấn đề tiếp "
   "cận đáng kể với một số nhóm thiểu số chủng tộc/sắc tộc (vd. Nam Á, Caribbean da "
   "đen, Phi) — vốn thường ngại tham vấn GP. Một LỘ TRÌNH SELF-REFERRAL do đó đã được "
   "phát triển cho các workshop để cho phép tiếp cận dễ dàng hơn. Điều này bao gồm "
   "quảng bá workshop rộng rãi tại các bối cảnh cộng đồng khác nhau, với THÔNG TIN LIÊN "
   "HỆ được làm nổi bật để cho phép tiếp cận trực tiếp.")
crit("Áp dụng VN: thay GP route → đăng ký qua tư vấn học đường, bằng kênh ẩn danh "
     "(app/web/QR code). Đặc biệt phù hợp văn hoá VN — HS Á Đông ngại nói trực tiếp "
     "với người lớn về vấn đề tâm lý.")

H2("Yếu tố 2 — PUBLICITY (P)")
en("Members of the public often struggle with aspects of mental health literacy, "
   "including not knowing what mental health problems are or if treatments will be "
   "effective, both of which are aspects of mental health literacy. Furthermore, "
   "self-reliance, or trying not to ask for help from others, is a common barrier.")
vn("Thành viên cộng đồng thường gặp khó khăn với các khía cạnh của mental health "
   "literacy — bao gồm không biết các vấn đề SKTT là gì hoặc các điều trị có hiệu quả "
   "không. Hơn nữa, self-reliance (tự dựa vào mình, không xin trợ giúp) là rào cản "
   "phổ biến.")

en("The aim of using detailed publicity was, therefore, to draw attention to the "
   "workshops and utilise a more attractive and engaging method to inform the public "
   "that more positive coping strategies were possible. Colourful and striking flyers "
   "were used to widely publicise workshops in libraries, pharmacies, community centres, "
   "and GP surgeries. Feedback about this approach was very positive.")
vn("Mục tiêu sử dụng publicity chi tiết là để THU HÚT SỰ CHÚ Ý đến workshop và sử dụng "
   "một phương pháp HẤP DẪN, THU HÚT để thông báo cho cộng đồng rằng các chiến lược đối "
   "phó tích cực hơn là khả dĩ. Các tờ rơi MÀU SẮC SỐNG ĐỘNG được sử dụng để quảng bá "
   "rộng rãi tại thư viện, hiệu thuốc, trung tâm cộng đồng và phòng khám GP. Phản hồi "
   "về cách tiếp cận này rất tích cực.")
crit("Áp dụng VN: poster + video trên mạng xã hội (Facebook, TikTok, Zalo) — kênh chính "
     "VTN VN dùng. Câu khẩu hiệu nên dùng ngôn ngữ tích cực kiểu DISCOVER: \"Bạn có "
     "muốn tự tin hơn?\" / \"Cách đối phó với áp lực thi cử\" — không dùng \"trầm cảm\" "
     "/ \"lo âu\".")

H2("Yếu tố 3 — ACCEPTABLE (AND ENGAGING) INTERVENTION (A)")
en("Acceptability of the treatment/intervention is another key barrier. Computerised CBT "
   "has not been found to be acceptable by many participants, especially when self-"
   "guided. The take-up rate is sometimes as low as 50%, with high rates of dropout. "
   "Indeed, a scoping review found that e-mental health treatment services were perceived "
   "as less helpful than traditional face-to-face interventions.")
vn("Acceptability của điều trị/can thiệp là một rào cản then chốt khác. CBT máy tính "
   "(computerised CBT) đã KHÔNG được coi là acceptable bởi nhiều người tham gia — đặc "
   "biệt khi tự hướng dẫn. Tỉ lệ tiếp nhận đôi khi thấp đến 50%, với tỉ lệ bỏ giữa "
   "chừng cao. Thực sự, một scoping review phát hiện rằng các dịch vụ e-mental health "
   "được cảm nhận là ÍT HỮU ÍCH HƠN so với can thiệp face-to-face truyền thống.")
crit("Đây là phát hiện QUAN TRỌNG cho VN: thay vì đầu tư app/online tự động (xu hướng "
     "đang được nhiều dự án Việt Nam theo đuổi), nên ưu tiên FACE-TO-FACE workshop. "
     "Tuy nhiên với thanh thiếu niên Gen Z VN có khả năng kỳ thị face-to-face nặng — "
     "cần feasibility test cụ thể trước khi quyết định.")

en("The stress workshops were designed to be run over one day and delivered in a group "
   "format to allow participants to share experiences without this process distracting "
   "from the psychoeducational focus. It was also designed around the concentration span "
   "of 20 min, with varied activities, including small and large group discussions, "
   "demonstrations of methods in role-plays by leaders, and individual exercises. The "
   "workshop was not specifically adapted to the needs of minority groups but was "
   "responsive to the different groups who attended. A colourful workbook covering the "
   "day's programme was provided to help participants sustain their progress.")
vn("Workshop stress được thiết kế chạy trong 1 ngày, theo định dạng NHÓM để cho phép "
   "người tham gia chia sẻ kinh nghiệm — mà không khiến quá trình này làm phân tâm khỏi "
   "trọng tâm tâm lý giáo dục. Cũng được thiết kế quanh KHẢ NĂNG TẬP TRUNG 20 PHÚT, "
   "với các hoạt động đa dạng — gồm thảo luận nhóm nhỏ và lớn, biểu diễn phương pháp "
   "qua role-play bởi leader, và bài tập cá nhân. Workshop KHÔNG được điều chỉnh cụ thể "
   "cho nhu cầu nhóm thiểu số, nhưng PHẢN ỨNG (responsive) với các nhóm khác nhau đã "
   "tham dự. Một sổ tay màu sắc bao quát chương trình ngày được cung cấp để giúp người "
   "tham gia duy trì tiến bộ.")

H2("Yếu tố 4 — CONVENIENCE (C)")
en("Services can be inconvenient when run during office hours between 9–5 pm and/or in "
   "formalized mental health settings. Weekly sessions during office hours may not suit "
   "everyone; participants may live far away from the therapy setting or may feel under "
   "pressure to get better quickly. Some patients also find it difficult to engage in "
   "lengthy psychological treatment and may prefer more intensive shorter treatments. "
   "There is surprisingly little literature around this topic.")
vn("Dịch vụ có thể KHÔNG THUẬN TIỆN khi chạy trong giờ hành chính 9-17h và/hoặc tại "
   "các cơ sở SKTT chính thức. Các buổi hàng tuần trong giờ hành chính có thể không "
   "phù hợp với mọi người; người tham gia có thể ở xa cơ sở trị liệu hoặc có thể cảm "
   "thấy áp lực phải hồi phục nhanh. Một số bệnh nhân cũng thấy khó engage với điều "
   "trị tâm lý kéo dài và có thể thích các điều trị NGẮN hơn nhưng INTENSIVE hơn. "
   "Đáng ngạc nhiên là có rất ít y văn về chủ đề này.")

en("Stress workshops were, therefore, run in a community setting, using leisure centres "
   "at the weekend to improve convenience and reduce possible stigma. The centre was "
   "conveniently situated for buses and car parking. While different, the move to more "
   "online interventions because of the pandemic is also important because it "
   "demonstrates the place of convenience.")
vn("Workshop stress do đó được tổ chức tại bối cảnh cộng đồng, sử dụng các TRUNG TÂM "
   "GIẢI TRÍ vào CUỐI TUẦN để cải thiện sự thuận tiện và giảm kỳ thị tiềm tàng. Trung "
   "tâm được đặt thuận tiện cho xe buýt và bãi đỗ ô tô. Mặc dù khác biệt, sự dịch "
   "chuyển sang nhiều can thiệp online do đại dịch cũng quan trọng vì chứng minh vị "
   "trí của convenience.")

en("Empirical outcomes of the stress workshops:\n"
   "(1) Uptake: 176 attended the information meeting.\n"
   "(2) Non-consultation: Just under half (41%) had not previously consulted their GPs.\n"
   "(3) Severity: In this study, participants' anxiety scores were above average "
   "(Spielberger trait scores 51.5), which is higher than the threshold for probable "
   "anxiety. A finding from a later study of Stress participants indicated 66% of stress "
   "workshop participants scored above the ICD psychiatric threshold.\n"
   "(4) Ethnicity: Data was not collected in the initial study but was collected in a "
   "later study in London, where 13.2% reported themselves as Black (9.4%) or Asian (3.8%).")
vn("Kết quả thực nghiệm của workshop stress:")
nr("• (1) Uptake: 176 người tham dự buổi thông tin.")
nr("• (2) Non-consultation: Chỉ dưới một nửa (41%) chưa từng tham vấn GP của họ.")
nr("• (3) Severity: Điểm lo âu Spielberger trait = 51,5 (trên ngưỡng \"probable anxiety\"); "
   "nghiên cứu sau cho thấy 66% người tham gia trên ngưỡng tâm thần ICD.")
nr("• (4) Ethnicity: 13,2% là Black (9,4%) hoặc Asian (3,8%) — thu hút được nhóm thiểu số.")
crit("Kết quả uptake và non-consultation 41% là một bằng chứng EVIDENCE tốt cho việc "
     "PLACES thu hút \"hard-to-reach\" group. Tuy nhiên thiếu CONTROL GROUP — không "
     "biết workshop \"thường\" sẽ đạt được bao nhiêu non-consulters. Đây là hạn chế "
     "thiết kế của 4 case studies.")

# ---------- 2.2 Depression workshops ----------
H2("2.2. Case 2 — Workshop Trầm cảm → \"Self-confidence Workshops\" (người lớn)")
en("When we used this initial model with pilot depression workshops, virtually all of "
   "the participants who attended had already been diagnosed as having depression and "
   "were being treated in primary or secondary care. We, therefore, added two new "
   "factors: using more lay non-diagnostic titles (L) and putting emphasis on the "
   "perceived effectiveness of the programme (E). These two new factors were added to "
   "Figure 1. We also changed the acceptability (A) of the programme by altering the "
   "description from being a depression programme to one on self-esteem/confidence.")
vn("Khi chúng tôi sử dụng mô hình ban đầu với pilot workshop trầm cảm, gần như TẤT CẢ "
   "người tham dự đã được chẩn đoán trầm cảm và đang được điều trị ở chăm sóc tuyến "
   "1/2. Vì vậy, chúng tôi đã THÊM HAI YẾU TỐ MỚI: sử dụng các tiêu đề LAY (L) — "
   "non-diagnostic — và nhấn mạnh hiệu quả CẢM NHẬN (E) của chương trình. Hai yếu tố "
   "mới này được thêm vào Hình 1. Chúng tôi cũng thay đổi acceptability (A) bằng cách "
   "đổi mô tả từ \"chương trình trầm cảm\" sang \"chương trình tự tin / lòng tự trọng\".")
crit("ĐIỂM MẤU CHỐT TRONG SỰ PHÁT TRIỂN MÔ HÌNH: từ 4 yếu tố ban đầu (S, P, A, C) "
     "→ 6 yếu tố sau khi gặp \"failure\" với depression workshop (chỉ thu hút người đã "
     "được chẩn đoán). Bằng chứng học từ THẤT BẠI này quan trọng — cho thấy mô hình "
     "evolves theo experience.")

H2("Yếu tố 5 — LAY NON-DIAGNOSTIC TITLES (L)")
en("Mental health literacy can have positive and negative effects. While it can be very "
   "helpful in providing an understanding of mental health problems, diagnostic labels "
   "can also act as a barrier to treatment and deter some people. Furthermore, a review "
   "of the effects of stigma and school interventions showed that the labelling of "
   "interventions compromised efforts to increase access to targeted school-based "
   "interventions. A number of studies show that an individual's perception of the "
   "problem is the most common barrier, with problems often perceived as 'social' (e.g., "
   "problems of living) rather than as 'medical' or 'psychological'.")
vn("Mental health literacy có thể có tác động tích cực VÀ tiêu cực. Trong khi có thể "
   "hữu ích trong việc cung cấp hiểu biết về vấn đề SKTT, NHÃN CHẨN ĐOÁN cũng có thể "
   "đóng vai trò RÀO CẢN tới điều trị và NGĂN một số người. Hơn nữa, một review về "
   "tác động của stigma và can thiệp trường học cho thấy việc DÁN NHÃN can thiệp đã "
   "làm giảm nỗ lực tăng tiếp cận các can thiệp targeted school-based. Một số nghiên "
   "cứu cho thấy CÁCH NHÌN CỦA CÁ NHÂN về vấn đề là rào cản phổ biến nhất — vấn đề "
   "thường được nhìn nhận là 'xã hội' (vd. vấn đề cuộc sống) hơn là 'y tế' hoặc 'tâm lý'.")

en("As 'depression' could be seen as a stigmatising diagnostic term, we, therefore, "
   "changed the title of the intervention from 'depression' to 'self-confidence'. This "
   "decision was also based on the close relationship between depression and self-esteem. "
   "The term 'self-confidence' was used rather than 'self-esteem' because the former is "
   "more often used colloquially and is more understandable to the public.")
vn("Vì 'trầm cảm' có thể được nhìn nhận như một thuật ngữ chẩn đoán mang tính kỳ thị, "
   "chúng tôi đã đổi tiêu đề can thiệp từ 'trầm cảm' sang 'tự tin'. Quyết định này cũng "
   "dựa trên mối quan hệ chặt chẽ giữa trầm cảm và lòng tự trọng. Thuật ngữ 'tự tin' "
   "(self-confidence) được sử dụng thay vì 'lòng tự trọng' (self-esteem) vì self-"
   "confidence được dùng phổ thông hơn và dễ hiểu hơn cho cộng đồng.")
crit("Đây là một QUYẾT ĐỊNH NGÔN NGỮ tinh tế: 'self-confidence' (tự tin) → đời thường, "
     "có khía cạnh tích cực; 'self-esteem' (lòng tự trọng) → vẫn mang tính kỹ thuật. "
     "Áp dụng VN: thay \"trầm cảm\" → \"tự tin\" / \"vượt khó\" / \"sống tích cực\". "
     "Lưu ý: cẩn thận với pop-psychology trends — \"self-confidence\" có thể bị "
     "hiểu thành \"thành công\"/\"giàu có\".")

H2("Yếu tố 6 — (PERCEIVED) EFFECTIVENESS OF INTERVENTION (E)")
en("When stigma was carefully examined, 'treatment stigma' (stigma specifically associated "
   "with seeking or receiving treatment for mental ill-health) was found to be one of the "
   "strongest predictors of low help-seeking. Among young people, the belief that the "
   "treatment would be effective was also shown to affect their willingness to seek help. "
   "Furthermore, this group perceived possible benefits as being more important than "
   "stigma-related factors. Perceived effectiveness of treatment was also found to "
   "influence seeking formal help among men.")
vn("Khi stigma được kiểm tra cẩn thận, 'treatment stigma' (kỳ thị cụ thể gắn với việc "
   "tìm hoặc nhận điều trị SKTT) được phát hiện là một trong những predictor MẠNH NHẤT "
   "của help-seeking thấp. Ở người trẻ, NIỀM TIN rằng điều trị sẽ HIỆU QUẢ cũng được "
   "chứng minh ảnh hưởng đến sự sẵn sàng tìm trợ giúp của họ. Hơn nữa, nhóm này nhìn "
   "nhận các LỢI ÍCH KHẢ DĨ là quan trọng hơn các yếu tố liên quan đến kỳ thị. Hiệu quả "
   "cảm nhận của điều trị cũng được phát hiện ảnh hưởng đến việc tìm trợ giúp chính "
   "thức ở NAM GIỚI.")
crit("Phát hiện \"perceived effectiveness > stigma\" rất quan trọng — đặc biệt cho VN "
     "nơi nam VTN ít tìm trợ giúp hơn nữ. Thông điệp publicity nên tập trung vào BENEFIT "
     "(\"tham gia giúp bạn ăn ngủ tốt hơn, tập trung học tốt hơn\") thay vì giảm stigma "
     "(\"không cần ngại\").")

en("In the publicity, it was therefore decided to market the potential effectiveness of "
   "the workshops more directly—\"Do you want to believe in yourself more? Handle times "
   "when things don't go your way? Be more effective in what you do? Put yourself down "
   "less often?\".")
vn("Trong publicity, do đó quyết định tiếp thị HIỆU QUẢ tiềm năng của workshop một cách "
   "trực tiếp hơn — \"Bạn có muốn tin tưởng vào bản thân nhiều hơn? Đối phó với những "
   "lúc mọi việc không như ý? Hiệu quả hơn trong việc bạn làm? Ít tự hạ mình hơn?\".")

H2("Yếu tố 3 mới — ACCEPTABILITY OF PROGRAMME (A) — đã điều chỉnh")
en("We also made changes to the programme to make it more acceptable (A). The content of "
   "the programme was changed to a CBT programme of self-confidence, described in "
   "Horrell, based on Fennell. The engaging format was again used with 20-min periods "
   "for the different methods, with interaction where possible. A colourful workbook "
   "was again provided to remind participants of the programme.")
vn("Chúng tôi cũng thay đổi chương trình để có acceptability hơn (A). Nội dung chương "
   "trình được đổi thành một chương trình CBT về tự tin — được mô tả trong Horrell, "
   "dựa trên Fennell. Định dạng thu hút được sử dụng lại với các giai đoạn 20 phút cho "
   "các phương pháp khác nhau, với tương tác khi có thể. Sổ tay màu sắc lại được cung "
   "cấp để nhắc nhở người tham gia về chương trình.")

vn("Các facilitator KHÁC (publicity P, self-referral S, convenience C) phần lớn được "
   "giữ nguyên.")

en("Empirical outcomes of depression workshops:\n"
   "Uptake: Changing the title from 'depression workshops' to 'self-confidence workshops' "
   "led to a marked increase in recruitment from 28 to 120 attendees.\n"
   "Non-consultation: After the change, the proportion of people who had not previously "
   "sought help increased from 9.8% to 39%. This suggests that depression is a term used "
   "by people already accessing mental health services.\n"
   "Severity of problems: 39% had not previously consulted GPs about depression. 72.6% "
   "(n = 106) of the workshop scored above the ICD psychiatric threshold.\n"
   "Equity: Self-confidence workshops did successfully engage ethnic minority groups, "
   "with 35.2% reporting themselves to be black (28.3%) which matches the ethnicity of "
   "the surrounding area, which was 25.9% black. This indicates more equitable access.")
vn("Kết quả thực nghiệm của workshop trầm cảm:")
nr("• Uptake: Đổi tiêu đề từ 'workshop trầm cảm' → 'workshop tự tin' dẫn đến tăng tuyển "
   "mộ ĐÁNG KỂ từ 28 → 120 người tham dự (tăng 4×).")
nr("• Non-consultation: Sau khi thay đổi, tỉ lệ người chưa từng tìm trợ giúp tăng từ "
   "9,8% → 39%. Điều này gợi ý rằng \"trầm cảm\" là thuật ngữ chỉ được dùng bởi những "
   "người đã tiếp cận dịch vụ SKTT rồi.")
nr("• Severity: 39% chưa từng tham vấn GP về trầm cảm; 72,6% (n=106) trên ngưỡng tâm "
   "thần ICD.")
nr("• Equity: 35,2% người tham gia là thiểu số sắc tộc (28,3% black) — KHỚP với thành "
   "phần dân cư xung quanh (25,9% black). → tiếp cận công bằng hơn.")
crit("BẰNG CHỨNG NGÔN NGỮ MẠNH: 28→120 attendees (tăng 4×) chỉ vì đổi tên. Tỉ lệ \"chưa "
     "tìm trợ giúp\" tăng từ 9,8% → 39% (gấp 4 lần) — nghĩa là tên \"trầm cảm\" đẩy lùi "
     "đúng nhóm cần can thiệp NHẤT (những người chưa được chăm sóc). Đây là dữ liệu "
     "convincing nhất trong cả paper PLACES. Áp dụng VN trực tiếp có thể được.")

en("Testing of 'PLACES' model: Given the success of the stress and self-confidence "
   "workshops in engaging participants, we decided to use and test this model with other "
   "age and clinical groups.")
vn("Kiểm thử mô hình 'PLACES': Cho thành công của các workshop stress và tự tin trong "
   "việc thu hút người tham gia, chúng tôi quyết định sử dụng và kiểm thử mô hình này "
   "với các nhóm tuổi và lâm sàng khác.")

# ---------- 2.3 Adolescent stress workshops ----------
H2("2.3. Case 3 — Workshop Stress cho VTN (DISCOVER)")
en("It was decided to adapt the adult version of the stress workshops for adolescents to "
   "offer early intervention and to test if this model could successfully engage "
   "adolescents who are traditionally reluctant to engage in services.")
vn("Quyết định ADAPT phiên bản người lớn của workshop stress cho VTN để cung cấp can "
   "thiệp sớm và kiểm thử xem mô hình có thể thành công trong việc thu hút VTN — vốn "
   "truyền thống ngại engage với dịch vụ — hay không.")

en("Extensive focus groups and interviews were conducted when adapting the workshops "
   "for adolescents. There were two major changes with this group: convenience (C) and "
   "acceptability (A). Lay (L) titles were also utilised.")
vn("Các focus group và interview rộng rãi đã được tiến hành khi adapt workshop cho VTN. "
   "Có HAI THAY ĐỔI LỚN với nhóm này: convenience (C) và acceptability (A). Tiêu đề "
   "Lay (L) cũng được sử dụng.")

H2("3a. CONVENIENCE — Setting & Timing (C)")
en("In the pilot study, attendance of the workshops was greater when these were run in "
   "schools compared to community settings such as libraries, youth clubs, or community "
   "centres. Schools were seen as safer, more convenient, and familiar environments to "
   "the adolescents than the other settings.")
vn("Trong nghiên cứu pilot, sự tham dự workshop CAO HƠN khi chúng được tổ chức tại "
   "TRƯỜNG so với các bối cảnh cộng đồng như thư viện, câu lạc bộ thanh niên, hoặc "
   "trung tâm cộng đồng. Trường được xem là an toàn hơn, thuận tiện hơn và quen thuộc "
   "hơn với VTN so với các bối cảnh khác.")
crit("Đây là PHÁT HIỆN đảo ngược với người lớn — người lớn thích ngoài-trường (leisure "
     "centre cuối tuần) để tránh stigma; VTN thì NGƯỢC LẠI thích trong-trường vì quen "
     "thuộc. Áp dụng VN: workshop tại trường + sau giờ học hoặc trong tiết HĐTN/sinh "
     "hoạt lớp là đúng hướng.")

H2("3b. ACCEPTABLE (AND ENGAGING) INTERVENTION (A)")
en("Strong preferences were expressed for more interactive and engaging content. It was "
   "also felt that a more individualised approach, focussing on goal planning, would be "
   "helpful in addition to the day-long workshop. A pre-workshop planning session, plus "
   "up to three phone calls after the workshop, were added.")
vn("Có ưu tiên MẠNH cho nội dung tương tác và thu hút hơn. Cũng cảm thấy rằng một cách "
   "tiếp cận CÁ NHÂN HOÁ hơn, tập trung vào lập kế hoạch mục tiêu, sẽ hữu ích — bổ sung "
   "cho workshop 1 ngày. Một buổi lên kế hoạch TRƯỚC workshop + tối đa 3 cuộc gọi điện "
   "thoại SAU workshop đã được thêm vào.")

H2("3c. LAY NON-DIAGNOSTIC TITLES (L)")
en("Adolescents suggested that the programme should be called the 'DISCOVER' workshop "
   "rather than just stress workshops. Some other facilitators were adapted, as outlined "
   "below.")
vn("VTN đề xuất chương trình nên được gọi là workshop 'DISCOVER' thay vì chỉ workshop "
   "stress. Các facilitator khác được điều chỉnh như mô tả dưới đây.")
crit("Tên DISCOVER là một CO-DESIGN OUTPUT — VTN tự đặt tên. Áp dụng VN: nên có buổi "
     "co-design với HS VN trước khi đặt tên chương trình; tên VN có thể là \"KHÁM PHÁ "
     "BẢN THÂN\" / \"VỮNG VÀNG\" / \"AN BÌNH\".")

H2("3d. SELF-REFERRAL (S) — Adolescent")
en("As self-reliance has been shown to be particularly relevant to adolescents, "
   "self-referral was seen as particularly valuable and emphasised.")
vn("Vì self-reliance được chứng minh đặc biệt liên quan đến VTN, self-referral được "
   "xem là ĐẶC BIỆT có giá trị và được nhấn mạnh.")

H2("3e. PUBLICITY (P) — Adolescent")
en("As well as paper publicity, popular social media platforms were also used.")
vn("Bên cạnh publicity giấy, các nền tảng MẠNG XÃ HỘI phổ biến cũng được sử dụng.")

en("Empirical outcomes of adolescent workshops:\n"
   "Uptake: Thirty-three participated in the pilot workshops and 155 attended the "
   "feasibility trial.\n"
   "Non-consultation: About 70% were non-consulters, 73.3% and 69.7% respectively.\n"
   "Equity: The rate for racial and ethnic students was high: 64.5% and 57.4%.\n"
   "Severity: Just under 50% scored above the threshold for anxiety problems on the RCADS-"
   "anxiety subscale, and just over 25% scored over the threshold for depression on the "
   "MFQ. In the pilot study, 74.2% scored above one or both of these clinical cut-offs.")
vn("Kết quả thực nghiệm của workshop VTN:")
nr("• Uptake: 33 trong pilot + 155 trong feasibility trial.")
nr("• Non-consultation: ~70% là non-consulters (73,3% pilot; 69,7% feasibility) — "
   "rất CAO ở VTN.")
nr("• Equity: Tỉ lệ HS chủng tộc/sắc tộc 64,5% (pilot) và 57,4% (feasibility) — RẤT CAO.")
nr("• Severity: ~50% trên ngưỡng RCADS-anxiety; ~25% trên ngưỡng MFQ depression; "
   "74,2% trên ít nhất 1 trong 2 ngưỡng (pilot).")
crit("Số liệu \"70% non-consulters\" + \"57-65% minority\" cho VTN là RẤT CONVINCING — "
     "PLACES thực sự thu hút được nhóm khó tiếp cận. Tuy nhiên cần lưu ý: pilot/"
     "feasibility KHÔNG có control group → không biết workshop \"thường\" sẽ thu hút "
     "non-consulters bao nhiêu %. RCT BESST 2024 (đã đọc) cho 80% non-consulters — "
     "consistent với pattern này.")

# ---------- 2.4 PND workshops ----------
H2("2.4. Case 4 — Workshop Trầm cảm Sau sinh (PND)")
en("Day-long PND workshops were developed as a result of recognising that, due to the "
   "small number of psychological therapists available, only a small proportion of "
   "mothers with PND (15%) were able to receive any evidence-based care. Long waitlists "
   "for psychotherapy, women's preferences for it over medication, a lack of time, and "
   "a reluctance to travel to regular appointments have been shown as substantial "
   "barriers. To increase engagement, all aspects of the 'PLACES' model were used but "
   "with some variations.")
vn("Workshop PND 1 ngày được phát triển sau khi nhận ra rằng — do số lượng ít các nhà "
   "trị liệu tâm lý có sẵn — chỉ một phần nhỏ (15%) các bà mẹ trầm cảm sau sinh có thể "
   "nhận chăm sóc dựa trên bằng chứng. Các rào cản đáng kể bao gồm: danh sách chờ dài "
   "cho psychotherapy, ưu tiên của phụ nữ với psychotherapy hơn thuốc, thiếu thời gian, "
   "và sự ngại di chuyển đến các cuộc hẹn thường xuyên. Để tăng engagement, TẤT CẢ "
   "các khía cạnh của mô hình 'PLACES' được sử dụng — với một số biến thể.")

H2("4a. PERCEIVED EFFECTIVENESS (E) — PND")
en("CBT techniques were used and capitalised on the fact that postpartum women prefer "
   "psychotherapies over medication. This highlighted that the workshop utilising "
   "non-pharmacological techniques was seen as important in increasing uptake, "
   "particularly among lactating mothers.")
vn("Các kỹ thuật CBT được sử dụng và tận dụng thực tế rằng phụ nữ sau sinh ưu tiên "
   "psychotherapy hơn thuốc. Điều này nhấn mạnh rằng workshop sử dụng kỹ thuật KHÔNG "
   "DÙNG THUỐC được xem là quan trọng trong tăng uptake — đặc biệt đối với các bà mẹ "
   "đang cho con bú.")

H2("4b. PUBLICITY AND SELF-REFERRAL (P và S) — PND")
en("Even though speciality psychiatric care in Canada normally requires a referral from "
   "a GP or another physician, the majority of the women self-referred (90–95%) to the "
   "workshop (and 80% were recruited via social media alone).")
vn("Mặc dù chăm sóc tâm thần chuyên khoa ở Canada thường cần giới thiệu từ GP hoặc bác "
   "sĩ khác, ĐA SỐ phụ nữ TỰ giới thiệu (90–95%) tới workshop (và 80% được tuyển qua "
   "mạng xã hội).")
crit("90-95% self-referral rate là CỰC KỲ CAO — bằng chứng mạnh PLACES vận hành ở Canada "
     "(không chỉ UK). Generalisability cross-country tốt. Áp dụng VN: kết hợp social "
     "media (Facebook, TikTok, Zalo) + self-referral là khả thi.")

H2("4c. CONVENIENCE (C) — PND")
en("The pilot face-to-face postnatal workshops were held in convenient locations on "
   "transit routes and in non-medical settings (community centres, libraries, etc.). "
   "Following consultation, these were run during the week as women with young children "
   "said they preferred weekday workshops, allowing more time with partners in the "
   "evenings or at weekends.")
vn("Workshop pilot face-to-face cho mẹ sau sinh được tổ chức tại các vị trí thuận tiện "
   "trên tuyến giao thông và trong các bối cảnh KHÔNG-y-tế (trung tâm cộng đồng, thư "
   "viện, v.v.). Sau khi tham vấn, các workshop này được chạy trong TUẦN — vì các bà "
   "mẹ có con nhỏ nói rằng họ thích workshop ngày tuần — cho phép có thêm thời gian "
   "với bạn đời vào tối hoặc cuối tuần.")

H2("4d. ACCEPTABILITY (A) — PND")
en("Extensive clinical work with mothers, collaboration with relevant public health, "
   "and community organizations helped to define workshop content, structure, and "
   "materials.")
vn("Công việc lâm sàng rộng rãi với các bà mẹ, sự hợp tác với y tế công cộng liên "
   "quan và các tổ chức cộng đồng đã giúp xác định nội dung, cấu trúc và tài liệu của "
   "workshop.")

en("Empirical outcomes of PND workshops:\n"
   "Uptake: Eighteen postnatally depressed women participated in the first pilot study. "
   "The online treatment involved 403 mothers.\n"
   "Non-consultation: Over half had not previously sought help; pilot 57% and online "
   "55%, respectively.\n"
   "Severity: All women had PND in both the pilot and online studies, scoring over 10 "
   "on Edinburgh Postnatal Scale (EPDS).")
vn("Kết quả thực nghiệm của workshop PND:")
nr("• Uptake: 18 phụ nữ trong pilot; 403 trong điều trị online.")
nr("• Non-consultation: Trên một nửa chưa từng tìm trợ giúp — pilot 57%, online 55%.")
nr("• Severity: Tất cả phụ nữ đều có PND (EPDS >10).")

# ---------- 3. Discussion ----------
H2("3. Thảo luận (Discussion)")
en("We have presented the implementation findings of the 'PLACES' model (Publicity, Lay, "
   "Acceptable, Convenient (Perceived) Effectiveness, Self-referral) of treatment "
   "engagement. The model incorporates both referral facilitators as well as intervention "
   "factors for improving engagement with mental health services/support. Four of the "
   "factors focus on the referral and publicising of the interventions (Publicity, Lay "
   "titles, (Perceived) Effectiveness, and Self-referral), and the other two factors "
   "relate to the intervention and its delivery (Acceptable and Convenient).")
vn("Chúng tôi đã trình bày các phát hiện triển khai của mô hình 'PLACES' (Publicity, "
   "Lay, Acceptable, Convenient, (Perceived) Effectiveness, Self-referral) của treatment "
   "engagement. Mô hình kết hợp cả các facilitator GIỚI THIỆU lẫn các yếu tố CAN THIỆP "
   "để cải thiện engagement với dịch vụ/hỗ trợ SKTT. BỐN yếu tố tập trung vào giới "
   "thiệu và publicising can thiệp (Publicity, Lay titles, Perceived Effectiveness, "
   "Self-referral), HAI yếu tố còn lại liên quan đến can thiệp và cách triển khai "
   "(Acceptable và Convenient).")
crit("PHÂN LOẠI 4+2 này là cấu trúc concept quan trọng. 4 yếu tố \"thượng nguồn\" — "
     "trước khi đến workshop. 2 yếu tố \"hạ nguồn\" — trong/sau workshop. Áp dụng VN: "
     "cần nỗ lực cả 2 phía: marketing/recruitment (P, L, E, S) + design workshop bản "
     "địa (A, C).")

en("This success of the model (with adaptations from the original stress workshop "
   "format) is that it has been shown to be effective in engaging people affected by "
   "depression, adolescents, those who have been traditionally difficult to engage, as "
   "well as mothers affected by PND who sometimes struggle to access to evidence-based "
   "psychological help.")
vn("Thành công của mô hình này (với các điều chỉnh từ định dạng workshop stress ban "
   "đầu) là nó đã được chứng minh hiệu quả trong việc thu hút những người bị trầm cảm, "
   "VTN, những người truyền thống khó thu hút, cũng như các bà mẹ bị PND vốn đôi khi "
   "khó tiếp cận trợ giúp tâm lý dựa trên bằng chứng.")

en("The 'PLACES' model is also very relevant to early intervention. It helps members of "
   "the public (whether adults or adolescents) think about their mental health needs, "
   "particularly when they are unsure about what to do next. This is particularly "
   "important because 'perceived need' has been found to be a key barrier to help-"
   "seeking. Given mental health services tend not to be publicised much, we believe that "
   "key elements have been the relevant 'social marketing' aspects: publicity (P) "
   "highlighting mental health problems in a more 'normal' way to lay people (L), "
   "highlighting the perceived effectiveness of the intervention (E), and setting up a "
   "self-referral system (S).")
vn("Mô hình 'PLACES' cũng rất phù hợp với CAN THIỆP SỚM. Nó giúp các thành viên cộng "
   "đồng (dù người lớn hay VTN) suy nghĩ về nhu cầu SKTT của mình — đặc biệt khi họ "
   "không chắc chắn phải làm gì tiếp theo. Điều này đặc biệt quan trọng vì 'perceived "
   "need' được phát hiện là rào cản then chốt cho help-seeking. Cho rằng dịch vụ SKTT "
   "thường KHÔNG được publicise nhiều, chúng tôi tin rằng các yếu tố then chốt là các "
   "khía cạnh 'social marketing' liên quan: publicity (P) làm nổi vấn đề SKTT theo cách "
   "'bình thường' hơn cho lay people (L), nhấn mạnh hiệu quả cảm nhận của can thiệp (E), "
   "và thiết lập hệ thống self-referral (S).")

en("Secondly, it tries to offer 'tailored' and effective psychosocial interventions that "
   "are acceptable (A) and convenient (C). Tailored interventions would 'make more "
   "sense' to those who would otherwise not use mental health services. Even though not "
   "widely researched, the acceptability of the intervention significantly affects "
   "take-up and dropout rates. A recent review shows the relevance of adapted "
   "interventions for racial and ethnic minorities.")
vn("Thứ hai, mô hình cố gắng cung cấp các can thiệp tâm lý xã hội 'TAILORED' và hiệu "
   "quả — có thể chấp nhận được (A) và thuận tiện (C). Can thiệp tailored sẽ 'có ý "
   "nghĩa hơn' với những người không sử dụng dịch vụ SKTT. Mặc dù không được nghiên "
   "cứu rộng rãi, acceptability của can thiệp ảnh hưởng đáng kể đến tỉ lệ tiếp nhận "
   "và bỏ giữa chừng. Một review gần đây cho thấy tính phù hợp của các can thiệp được "
   "điều chỉnh cho nhóm thiểu số chủng tộc/sắc tộc.")

en("Convenience is another important facilitator within the PLACES model. Whereas this "
   "facilitator has been informed by and tested during pre-COVID-19 case studies, the "
   "huge drive towards the online delivery of mental health treatment to ensure "
   "continuity of care, and its possible effect on the implementation of our model, "
   "needs to be considered. For some members of the public, the provision of online "
   "therapy has been a convenient change, whereas other people may have been excluded "
   "due to limited technological literacy, lack of resources, or preference for face-"
   "to-face sessions.")
vn("Convenience là một facilitator quan trọng khác trong mô hình PLACES. Dù facilitator "
   "này đã được informed và kiểm thử trong các case studies trước COVID-19, sự thúc "
   "đẩy lớn hướng tới cung cấp điều trị SKTT online để đảm bảo tính liên tục của chăm "
   "sóc, và tác động khả dĩ của nó lên việc triển khai mô hình của chúng tôi, cần được "
   "xem xét. Với một số thành viên cộng đồng, cung cấp therapy online là thay đổi "
   "thuận tiện; trong khi những người khác có thể bị loại trừ do hạn chế literacy công "
   "nghệ, thiếu nguồn lực, hoặc ưu tiên buổi face-to-face.")

en("With self-referral, the 'PLACES' model has successfully engaged people from diverse "
   "backgrounds as well as those with serious problems and has led to the Improving "
   "Access to Psychological Treatments (IAPT) service in the UK to adopt this method. "
   "Furthermore, self-referral was found to attract a more equitable group from ethnic "
   "minorities and employment status compared to GP referrals. A similar picture was "
   "found with a self-referral system in the military which was also found to have high "
   "rates of non-consulters (69%) as well as those with serious problems (72%).")
vn("Với self-referral, mô hình 'PLACES' đã thành công trong việc thu hút người từ các "
   "background đa dạng cũng như những người có vấn đề nặng — và đã dẫn đến dịch vụ "
   "Improving Access to Psychological Treatments (IAPT) ở UK ÁP DỤNG phương pháp này. "
   "Hơn nữa, self-referral được phát hiện thu hút một nhóm CÔNG BẰNG HƠN từ thiểu số "
   "sắc tộc và tình trạng việc làm so với GP referrals. Một bức tranh tương tự được "
   "tìm thấy với hệ thống self-referral trong quân đội — cũng có tỉ lệ cao non-consulters "
   "(69%) cũng như người có vấn đề nặng (72%).")
crit("THÔNG TIN QUAN TRỌNG: PLACES self-referral đã được IAPT (UK national service) "
     "ÁP DỤNG — đây là validation cấp HỆ THỐNG, không chỉ trong studies. Áp dụng VN: "
     "có thể đề xuất với Bộ Y tế hoặc Bộ GD-ĐT VN tích hợp self-referral vào tư vấn "
     "học đường.")

en("The proposed model is unique for several reasons. Firstly, it fills an important "
   "gap in the literature about the help-seeking behaviour of the general public. Two "
   "major reviews have concluded that effective methods that change the help-seeking "
   "behaviour of the general public have not been identified. The 'PLACES' model focuses "
   "on changing help-seeking behaviour rather than shifting attitudes alone. Secondly, "
   "it uniquely combines facilitators with an engaging intervention for the general "
   "public. It describes more of a 'bridge' that a person with a mental health problem "
   "would cross in order to reach services that up to now have been avoided or are "
   "unknown. This model is, therefore, about trying to make the services more "
   "'reachable'. Thirdly, it is based on the wider literature, as well as empirical work "
   "conducted in developing the intervention of large-scale one-day psychoeducational "
   "workshops in the community. Finally, it aims to be clinically relevant so that "
   "services can apply these 'bridging' principles in the real world. Clinicians and "
   "service managers could find the changes in this model, which are reasonably small "
   "and easy to implement, useful if they are trying to engage particular groups, such "
   "as specific racial and ethnic minorities.")
vn("Mô hình đề xuất này UNIQUE vì nhiều lý do. THỨ NHẤT, nó lấp một khoảng trống quan "
   "trọng trong y văn về hành vi help-seeking của cộng đồng. Hai review lớn đã kết "
   "luận rằng các phương pháp hiệu quả thay đổi hành vi help-seeking của cộng đồng "
   "CHƯA ĐƯỢC NHẬN DIỆN. Mô hình 'PLACES' tập trung vào THAY ĐỔI HÀNH VI help-seeking "
   "thay vì chỉ chuyển đổi thái độ. THỨ HAI, mô hình kết hợp ĐỘC ĐÁO các facilitator "
   "với một can thiệp thu hút cho cộng đồng. Mô hình mô tả nhiều hơn một 'CÂY CẦU' mà "
   "người có vấn đề SKTT phải vượt qua để tiếp cận các dịch vụ — vốn cho đến nay đã bị "
   "tránh hoặc không được biết đến. Do đó mô hình này về việc cố gắng làm cho dịch vụ "
   "REACHABLE hơn. THỨ BA, dựa trên y văn rộng cũng như công trình thực nghiệm trong "
   "phát triển can thiệp workshop tâm lý giáo dục 1 ngày quy mô lớn trong cộng đồng. "
   "Cuối cùng, mô hình hướng tới có liên quan lâm sàng — để các dịch vụ có thể áp dụng "
   "các nguyên lý 'cây cầu' này trong thực tế. Các nhà lâm sàng và quản lý dịch vụ có "
   "thể thấy các thay đổi trong mô hình này — vốn KHÁ NHỎ và DỄ TRIỂN KHAI — hữu ích "
   "nếu họ đang cố gắng thu hút các nhóm cụ thể như các nhóm thiểu số chủng tộc/sắc tộc.")

en("We believe this model has the potential to reduce the prevalence of mental health "
   "problems by increasing the number of people in the community who are able to access "
   "evidence-based treatments, and with that, reduce the chronicity of problems. "
   "Additionally, it could have financial benefits as well as improved recruitment to "
   "research studies.")
vn("Chúng tôi tin rằng mô hình này có TIỀM NĂNG GIẢM TỈ LỆ vấn đề SKTT bằng cách tăng "
   "số người trong cộng đồng tiếp cận điều trị dựa trên bằng chứng, và với đó, giảm "
   "tính mạn tính của vấn đề. Bổ sung, mô hình có thể có lợi ích tài chính cũng như "
   "cải thiện tuyển mẫu cho các nghiên cứu.")

en("Several help-seeking models exist. We have mentioned the process model by Gask, "
   "which covers engagement/promotion and intervention aspects. Other well-known models "
   "of help-seeking, such as the 'Behavioural Model of Health Services use' and the "
   "Theory of Planned Behaviour, focus on the attitudinal aspects prior to intervention.")
vn("Có nhiều mô hình help-seeking. Chúng tôi đã đề cập mô hình tiến trình của Gask — "
   "bao quát các khía cạnh engagement/promotion và intervention. Các mô hình "
   "help-seeking nổi tiếng khác — như 'Behavioural Model of Health Services use' "
   "(Andersen) và Theory of Planned Behaviour (Ajzen) — tập trung vào các khía cạnh "
   "thái độ TRƯỚC khi can thiệp.")

# ---------- 4. Clinical Implications ----------
H2("4. Hệ quả lâm sàng (Clinical Implications)")
en("Clinical services may fear being overwhelmed by demand if services are better "
   "publicised. However, this model complements existing services. A stepped care model "
   "already occurs with IAPT, or a staging process for early intervention is being "
   "developed, as relatively few people would need intensive treatment. Those with no or "
   "mild symptoms could be signposted to self-help material or relevant websites, and "
   "those with more severe problems can be directed to more intensive services.")
vn("Các dịch vụ lâm sàng có thể LO NGẠI bị quá tải nhu cầu nếu dịch vụ được publicise "
   "tốt hơn. Tuy nhiên, mô hình này BỔ SUNG cho dịch vụ hiện có. Mô hình stepped care "
   "đã có ở IAPT, hoặc quy trình staging cho can thiệp sớm đang được phát triển — vì "
   "tương đối ít người cần điều trị intensive. Những người không/triệu chứng nhẹ có "
   "thể được hướng dẫn đến tài liệu tự giúp hoặc website liên quan, những người vấn "
   "đề nặng hơn có thể được hướng tới dịch vụ intensive hơn.")
crit("Cảnh báo \"overwhelm services\" là rất thực tế. Áp dụng VN: trước khi triển khai "
     "PLACES quy mô lớn, cần có CAPACITY PLANNING cho tư vấn học đường — không thể tăng "
     "demand mà không tăng supply. Nếu không sẽ dẫn đến disappointment + chán nản hệ "
     "thống.")

H2("Hệ quả Nghiên cứu (Research Implications)")
en("Identifying how 'PLACES' factors could be applied to different clinical or "
   "demographic groups—and identifying the more relevant facilitators for these groups—"
   "would be valuable. For instance, normalisation has been found to be particularly "
   "important for boys/students in describing mental health problems, as they have fear "
   "of being seen as 'weak'. Other groups where the 'PLACES' model could be applied may "
   "be south Asians and Caribbean men. For each of these groups, different methods for "
   "marketing and/or tailoring interventions could be compiled and systematically tested "
   "to see if the different methods result in better engagement. Treatment engagement "
   "scales can also be developed.")
vn("Xác định cách các yếu tố 'PLACES' có thể được áp dụng cho các nhóm lâm sàng/nhân "
   "khẩu khác nhau — và xác định các facilitator liên quan hơn cho các nhóm này — sẽ "
   "có giá trị. Ví dụ, NORMALISATION (chuẩn hoá / bình thường hoá) được phát hiện đặc "
   "biệt quan trọng cho NAM SINH trong việc mô tả vấn đề SKTT — vì họ sợ bị xem là "
   "'yếu đuối'. Các nhóm khác mà mô hình 'PLACES' có thể được áp dụng là Nam Á và Nam "
   "giới Caribbean. Đối với mỗi nhóm này, các phương pháp marketing và/hoặc tailoring "
   "can thiệp khác nhau có thể được biên soạn và test có hệ thống để xem liệu các "
   "phương pháp khác nhau có dẫn đến engagement tốt hơn hay không. Các thước đo "
   "treatment engagement cũng có thể được phát triển.")

en("In order to make the best use of interventions that are clinically effective for "
   "people who need them, we should reduce the \"treatment engagement gap\" between "
   "treatments and seeking help. If the treatment engagement factors outlined in the "
   "'PLACES' model are shown to be helpful with different groups, then these methods "
   "could be very helpful in reducing the prevalence of mental health problems and with "
   "it related benefits.")
vn("Để sử dụng tốt nhất các can thiệp hiệu quả lâm sàng cho những người cần chúng, "
   "chúng ta nên GIẢM \"khoảng trống treatment engagement\" giữa điều trị và tìm trợ "
   "giúp. Nếu các yếu tố treatment engagement được nêu trong mô hình 'PLACES' được "
   "chứng minh hữu ích với các nhóm khác nhau, thì các phương pháp này có thể rất "
   "hữu ích trong việc giảm tỉ lệ vấn đề SKTT và các lợi ích liên quan.")

# =====================================================================
# PHẦN 2 — BẢNG TỔNG HỢP
# =====================================================================
H1("PHẦN 2 — BẢNG TỔNG HỢP")

H2("Bảng A. 6 thành phần của mô hình PLACES — Định nghĩa, bằng chứng, ứng dụng")
tbl(['Yếu tố', 'Tiếng Anh', 'Bản chất', 'Bằng chứng từ 4 case studies', 'Ứng dụng cho VN'],
    [
        ['P', 'Publicity',
         'Quảng bá tích cực, qua nhiều kênh, ngôn ngữ tích cực',
         'Tờ rơi màu sắc + social media (PND 80% qua MXH); flyer thư viện/hiệu thuốc',
         'Facebook + TikTok + Zalo; poster trường + tờ rơi GVCN; '
         'thông điệp tích cực ("tự tin"/"vững vàng")'],
        ['L', 'Lay (non-diagnostic) titles',
         'Đặt tên phi chẩn đoán, đời thường',
         'Đổi "depression" → "self-confidence" → uptake 28→120 (4×); '
         'non-consulter 9.8%→39%',
         'Thay "lo âu/trầm cảm" → "căng thẳng học tập"/"phòng kỹ năng sống"/'
         '"khám phá bản thân"'],
        ['A', 'Acceptable (& engaging) intervention',
         'Workshop tương tác, group format, 20-min activities',
         'Sổ tay màu sắc, role-play, video clips; HS thích interactive (DISCOVER)',
         'Adapt văn hoá Á Đông: face-saving, family-respect; module ngắn 20 phút; '
         'group nhỏ 8-12 HS'],
        ['C', 'Convenient location & timing',
         'Bối cảnh không-y-tế, gần, giờ phù hợp',
         'Adults thích leisure centre cuối tuần; VTN thích trường + sau giờ học; '
         'PND ngày tuần',
         'Tại trường + tiết HĐTN/sinh hoạt lớp; cuối tuần cho mở rộng cộng đồng'],
        ['E', '(Perceived) Effectiveness',
         'Thông điệp marketing chuyên về benefit, không về stigma',
         '"Do you want to believe in yourself more?"; bằng chứng nam giới response '
         'với perceived effectiveness > stigma reduction',
         'Slogan VN: "Học giỏi hơn / ngủ ngon hơn / tự tin trước thi cử"; '
         'không nói "giảm stigma"'],
        ['S', 'Self-referral',
         'HS tự đăng ký, không qua GP/GVCN/phụ huynh',
         'Adults: 41% non-consulters; PND: 90-95% self-refer; VTN: 70% non-consulters',
         'App/web ẩn danh + QR code trong giờ chào cờ; tránh kênh GVCN '
         'để giảm stigma'],
    ], [1.0, 3.0, 3.5, 4.5, 4.5])

H2("Bảng B. Empirical results — 4 case studies")
tbl(['Case study', 'N', 'Uptake', '% Non-consulters', '% Severity', '% Minority equity'],
    [
        ['1. Stress workshop (adults)', '176',
         '176 attended info', '41%',
         '66% above ICD threshold', '13.2% (Black 9.4% + Asian 3.8%)'],
        ['2. Self-confidence workshop (adults, formerly depression)',
         '~120 (sau đổi tên từ 28)',
         '4× tăng sau đổi tên', '39% (tăng từ 9.8%)',
         '72.6% above ICD (n=106)',
         '35.2% minority (28.3% black) — match local 25.9%'],
        ['3. Adolescent stress workshop (DISCOVER)',
         '33 pilot + 155 feasibility',
         'Cao (155 trong 1 trial)',
         '~70% (73.3%/69.7%)',
         '~50% RCADS-anxiety, 25% MFQ depression, 74.2% trên ≥1 cutoff',
         '64.5%/57.4% racial/ethnic'],
        ['4. PND workshop (Canada)',
         '18 pilot + 403 online',
         '90-95% self-referred (80% qua social media)',
         '~55-57%',
         '100% PND (EPDS >10)',
         'Không báo cáo riêng'],
    ], [3.0, 2.0, 2.5, 2.5, 2.8, 3.5])

H2("Bảng C. So sánh PLACES với 3 mô hình help-seeking khác")
tbl(['Mô hình', 'Tác giả/Năm', 'Trọng tâm', 'Áp dụng cho hành vi behaviour?', 'Ưu/Nhược'],
    [
        ['PLACES (6 yếu tố)', 'Brown JSL et al. 2022',
         'Treatment engagement (cầu nối user → service)',
         'CÓ — đo behaviour (uptake, non-consultation rate)',
         'Ưu: practical, đã được IAPT UK adopt; Nhược: case studies '
         'không có control group, level 4 evidence'],
        ['Behavioural Model of Health Services Use', 'Andersen 1995 (Andersen-Newman 1973)',
         'Predisposing + Enabling + Need factors',
         'CÓ — nhưng chủ yếu cấu trúc xã hội',
         'Ưu: comprehensive (cá nhân + hệ thống); Nhược: phức tạp, khó operationalise'],
        ['Theory of Planned Behaviour (TPB)', 'Ajzen 1991',
         'Attitude + Subjective norm + Perceived behavioural control → Intention',
         'GIÁN TIẾP — qua intention, không phải behaviour trực tiếp',
         'Ưu: chuẩn academic, được test rộng; Nhược: '
         'attitude-behaviour gap kinh điển'],
        ['Health Belief Model (HBM)', 'Rosenstock 1974',
         'Perceived susceptibility + severity + benefits + barriers + cues',
         'GIÁN TIẾP — qua perceived components',
         'Ưu: lâu đời, intuitive; Nhược: bỏ qua emotion + social factors'],
        ['Gask conceptual model', 'Gask et al. 2012',
         'Candidacy → Navigation → Appearance (community engagement)',
         'CÓ — patient-centred process model',
         'Ưu: process-based, cover uncertainty; Nhược: mô tả không '
         'thành công cụ can thiệp'],
    ], [3.0, 3.0, 4.0, 3.5, 4.5])

H2("Bảng D. Adapt PLACES cho Việt Nam — đề xuất chi tiết")
tbl(['Yếu tố', 'Bản gốc UK', 'Adapt cho VN', 'Cần thận trọng / Validate'],
    [
        ['P (Publicity)',
         'Tờ rơi + flyer thư viện/GP',
         'Facebook + TikTok + Zalo; poster lớp; thông điệp qua phụ huynh',
         'MXH có thể overload; cần A/B test thông điệp; '
         'phụ huynh literacy thấp (5.1%) khó truyền tin tốt'],
        ['L (Lay titles)',
         '"Self-confidence workshop"',
         '"Khám phá bản thân"/"Vững vàng"/"Lớp kỹ năng đối phó stress"',
         'Cần co-design với HS VN để chọn tên; '
         'tránh từ trendy ("self-help" có thể bị hiểu là pop-psych)'],
        ['A (Acceptable)',
         'Group 30 người, 20-min modules',
         'Group 8-12 HS (lớp nhỏ hơn UK); module 15-20 phút; '
         'thêm role-play tình huống VN (áp lực thi đại học, gia đình)',
         'Văn hoá face-saving — HS có thể ngại nói trong nhóm; '
         'cân nhắc one-to-one option'],
        ['C (Convenient)',
         'Cuối tuần / leisure centre (adults); '
         'tại trường giờ học (VTN)',
         'Tại trường + tiết HĐTN/sinh hoạt; '
         'option ngoài giờ (cuối tuần) cho HS muốn ẩn danh hơn',
         'Tiết HĐTN có thể bị GV chính khoá lấy mất; '
         'cần policy bảo vệ slot can thiệp'],
        ['E (Perceived Effectiveness)',
         'Slogan benefit-driven',
         '"Học tốt hơn — ngủ ngon hơn — tự tin trước thi"; '
         'kèm con số benefit nếu có RCT VN',
         'Chưa có RCT VN → "perceived" hiện chỉ dựa trên '
         'evidence UK; cần feasibility VN'],
        ['S (Self-referral)',
         'Đăng ký trực tiếp qua publicity',
         'App/web ẩn danh + QR code; '
         'option qua GVCN cho HS không tự tin',
         'GVCN có thể gây stigma — cần tuyệt đối ẩn danh; '
         'phụ huynh có thể phát hiện/cản trở'],
    ], [2.5, 3.5, 5.5, 5.5])

# =====================================================================
# PHẦN 3 — PHẢN BIỆN TỔNG QUAN
# =====================================================================
H1("PHẦN 3 — PHẢN BIỆN TỔNG QUAN")

H2("3.1. Điểm mạnh của bài và mô hình PLACES")
nr("• Lấp một khoảng trống lý thuyết quan trọng — chuyển trọng tâm từ ATTITUDE sang "
   "BEHAVIOUR.")
nr("• Mô hình SIMPLE và OPERATIONAL — 6 yếu tố dễ nhớ, dễ áp dụng.")
nr("• Đã được VALIDATION cấp HỆ THỐNG: IAPT (Improving Access to Psychological "
   "Treatments) UK — dịch vụ quốc gia — đã áp dụng self-referral.")
nr("• 4 case studies COVER nhiều populations (adults stress/depression, VTN, mẹ PND) "
   "— gợi ý generalisability.")
nr("• PHÁT TRIỂN TIẾN HOÁ qua thực tiễn (4 → 6 yếu tố sau khi gặp \"failure\" với "
   "depression workshop) — cho thấy mô hình LIVING, không cố định.")
nr("• Ngôn ngữ accessible — nhà thực hành có thể đọc và áp dụng ngay.")

H2("3.2. Điểm yếu / Hạn chế của mô hình và bài")
crit("(1) LOẠI BẰNG CHỨNG: 4 case studies = level-4 evidence (Oxford CEBM 2011) — "
     "case series không có control group song song. Không đủ để kết luận causal "
     "(\"PLACES gây ra increased uptake\"). Tới BESST 2024 mới có RCT cho một phần "
     "PLACES (Self-referral) — nhưng PLACES TOÀN BỘ 6 yếu tố CHƯA TỪNG được test "
     "trong RCT độc lập.")
crit("(2) SELECTION BIAS: 4 case studies đều của CHÍNH NHÓM TÁC GIẢ — risk of "
     "publication bias (chỉ public các case thành công) và investigator bias.")
crit("(3) CONFOUNDING: khi đổi tên workshop từ 'depression' → 'self-confidence', uptake "
     "tăng 4× — nhưng cũng có thể do thời điểm khác (campaign khác, năm khác, kinh "
     "phí khác). Không có A/B test song song.")
crit("(4) OVERFITTING TO UK CONTEXT: 3/4 case studies ở UK; chỉ PND ở Canada. "
     "Generalisability sang LMIC (Việt Nam, Châu Phi) chưa được test. Đặc biệt yếu tố "
     "C (convenience, leisure centre, transit routes) phụ thuộc INFRASTRUCTURE đô thị "
     "phương Tây.")
crit("(5) BỎ QUA STRUCTURAL BARRIERS: PLACES tập trung individual + service level — "
     "không cover policy, funding, training infrastructure. Áp dụng VN cần thêm structural "
     "layer (policy MOET, funding, đào tạo tư vấn học đường).")
crit("(6) THIẾU OPERATIONAL DEFINITION cho từng yếu tố: \"Acceptable\" cụ thể là gì? "
     "Đo bằng thang nào? Bao nhiêu engagement-rate là \"acceptable\"? Bài không cung "
     "cấp criteria cụ thể → khó replicate.")
crit("(7) THIẾU NEGATIVE CASES: bài chỉ trình bày 4 case THÀNH CÔNG. Có case nào "
     "PLACES KHÔNG work? Tại sao? Bỏ sót thông tin này làm mô hình thiếu cân bằng.")
crit("(8) CHƯA RIGOROUS ENGAGEMENT MEASURES: tác giả tự thừa nhận \"treatment engagement "
     "scales can also be developed\" — tức là HIỆN CHƯA CÓ thước đo engagement chuẩn "
     "hoá. Đo bằng \"uptake\" + \"non-consulter %\" + \"minority %\" không đủ.")

H2("3.3. Phù hợp với bối cảnh Việt Nam")
nr("• PLACES có TIỀM NĂNG MẠNH cho VN, đặc biệt 4 yếu tố:")
nr("    – L (Lay titles): bằng chứng rõ ràng (28→120 attendees) — áp dụng trực tiếp "
   "thay tên \"can thiệp lo âu\" → \"workshop căng thẳng học tập\" / \"khám phá bản "
   "thân\"")
nr("    – S (Self-referral): essential cho VN nơi VTN ngại tham vấn người lớn — "
   "cần app/web ẩn danh")
nr("    – E (Perceived Effectiveness): thông điệp benefit-driven phù hợp văn hoá "
   "VN \"học tốt - sống tốt\"")
nr("    – C (Convenient): trường + giờ HĐTN — VN đã có cấu trúc HĐTN trong chương "
   "trình GD-PT 2018")
nr("• 2 yếu tố cần ADAPT CẨN THẬN cho VN:")
nr("    – P (Publicity): kênh khác UK (MXH thay flyer); literacy phụ huynh thấp 5,1% "
   "nên thông điệp phải tới HS trực tiếp + qua peer influencer")
nr("    – A (Acceptable): cần co-design với HS VN — văn hoá face-saving khác UK")
nr("• KHUYẾN NGHỊ TRIỂN KHAI VN:")
nr("    – Bước 1: Co-design với 30-50 HS VN (focus groups + interviews) để adapt 6 yếu "
   "tố")
nr("    – Bước 2: Feasibility study (n=100-200, 1-2 trường) đo uptake + acceptability")
nr("    – Bước 3: Pilot RCT (n=400-600, 5-10 trường) đo cả engagement + clinical outcome")
nr("    – Bước 4: Scale up nếu pilot positive — kết hợp policy advocacy với MOET/MOH")

H2("3.4. Đối chiếu với corpus 35+ bài của dự án")
tbl(['Bài/RAG ID', 'Tên / Tác giả', 'Liên hệ với mô hình PLACES'],
    [
        ['QT042 / B5', 'Brown & Carter 2025 Editorial',
         'Tổng hợp lập trường tác giả; áp dụng PLACES trong BESST trial — tổng hợp tốt'],
        ['BESST trial (Brown 2024)', 'Lancet Psychiatry RCT',
         'Test PLACES Self-referral component trong school context — '
         'd=−0.17 chung, d=−0.52 subgroup elevated depression'],
        ['QT08 (Wen 2020)', 'LPA lo âu nông thôn TQ',
         'Hỗ trợ trường OR=0,562 (bảo vệ) — củng cố vai trò trường (yếu tố C của PLACES)'],
        ['VN030 (Happy House)', 'VN universal school',
         'd=0,11 (rất nhỏ) — diluting effect khi không áp dụng targeted self-referral '
         'theo PLACES'],
        ['INSIGHT_05 (corpus)', 'Yếu tố nguy cơ + bảo vệ VN',
         'Hỗ trợ trường = bảo vệ (Wen 2020); cha mẹ tham gia AOR=0,75 — cần adapt '
         'P + S cho cha mẹ ở VN'],
        ['QT29 (Li 2025)', 'NMA can thiệp lo âu trẻ em 12 nước',
         'CBT hạng 2 SUCRA (sau ACT) — content của workshop PLACES (CBT-based) '
         'phù hợp evidence'],
        ['B6 (Resilience)', 'School resilience review',
         'Resilience programs cũng dùng workshop format — nhưng chưa rigorous '
         'như PLACES về engagement'],
    ], [3.0, 4.0, 10.0])

# =====================================================================
# PHẦN 4 — THAM KHẢO
# =====================================================================
H1("PHẦN 4 — THAM KHẢO ĐẦY ĐỦ (61 references gốc + nguồn xác minh)")
note("Bài Brown 2022 PLACES có 61 references. Em liệt kê đầy đủ các ref quan trọng + "
     "nguồn xác minh ngoài. Mỗi ref giữ format APA + DOI khi có. Trang ref gốc: "
     "trang 9-12 của PDF.")

H2("4.1. Reference chính của bài (đã trích đầy đủ — chỉ list các ref có DOI)")
nr("1. McManus S, Bebbington PE, Jenkins R, Brugha T (2016). Mental Health and Wellbeing "
   "in England: Adult Psychiatric Morbidity Survey 2014. NHS Digital, Leeds. "
   "ISBN 978-1-78386-738-6", size=10)
nr("2. Andrews G, Sanderson K, Slade T, Issakidis C (2000). Why does the burden of "
   "disease persist? Bull World Health Organ, 78:446–454. PMID 10885163", size=10)
nr("3. Brown JSL, Murphy C, Kelly J, Goldsmith K (2019). How can we successfully recruit "
   "depressed people? CLASSIC trial. Trials, 20:131. doi:10.1186/s13063-018-3033-5 — "
   "PMID 30760332", size=10)
nr("4. Gulliver A, Griffiths KM, Christensen H, Brewer JL (2012). A SR of help-seeking "
   "interventions for depression, anxiety and general distress. BMC Psychiatry, 12:81. "
   "doi:10.1186/1471-244X-12-81 — PMID 22812691", size=10)
nr("5. Xu Z, Huang F, Koesters M, Staiger T, Becker T, Thornicroft G, Ruesch N (2018). "
   "Effectiveness of interventions to promote help-seeking: SR & meta-analysis. Psychol "
   "Med, 48:2658–2667. doi:10.1017/S0033291718001265 — PMID 29852885", size=10)
nr("6. Kessler RC et al. (2005). Lifetime prevalence and age-of-onset of DSM-IV "
   "disorders. Arch Gen Psychiatry, 62:593–602. doi:10.1001/archpsyc.62.6.593 — "
   "PMID 15939837", size=10)
nr("7. Colizzi M, Lasalvia A, Ruggeri M (2020). Prevention and early intervention in "
   "youth mental health. Int J Ment Health Syst, 14:23. doi:10.1186/s13033-020-00356-9 — "
   "PMID 32226481", size=10)
nr("8. Gask L et al. (2012). Improving access to psychosocial interventions for common "
   "mental health problems in the UK. BMC Health Serv Res, 12:249. "
   "doi:10.1186/1472-6963-12-249 — PMID 22876746", size=10)
nr("9. Clement S et al. (2014). What is the impact of mental health-related stigma on "
   "help-seeking? SR. Psychol Med, 45:11–27. doi:10.1017/S0033291714000129 — "
   "PMID 24569086", size=10)
nr("10. Mojtabai R et al. (2010). Barriers to mental health treatment: NCS-R. Psychol "
    "Med, 41:1751–1761. doi:10.1017/S0033291710002291 — PMID 21134315", size=10)
nr("11. Bonabi H et al. (2016). Mental Health Literacy, Attitudes, and Perceived Need "
    "as Predictors of MHS Use. J Nerv Ment Dis, 204:321–324. "
    "doi:10.1097/NMD.0000000000000488 — PMID 26894319", size=10)
nr("12. Codony M et al. (2009). Perceived Need for MHS and Service Use among Adults in "
    "Western Europe: ESEMeD. Psychiatr Serv, 60:1051–1058. "
    "doi:10.1176/ps.2009.60.8.1051 — PMID 19648192", size=10)
nr("13. Rafferty LA, Wessely S, Stevelink SA, Greenberg N (2019). Journey to "
    "professional MH support: military veterans. Eur J Psychotraumatol, 10:1700613. "
    "doi:10.1080/20008198.2019.1700613 — PMID 31908738", size=10)
nr("14. Gulliver A, Griffiths KM, Christensen H (2010). Perceived barriers and "
    "facilitators in young people: SR. BMC Psychiatry, 10:113. "
    "doi:10.1186/1471-244X-10-113 — PMID 21192795", size=10)
nr("15. Jorm AF (2000). Mental health literacy. Br J Psychiatry, 177:396–401. "
    "doi:10.1192/bjp.177.5.396 — PMID 11059991", size=10)
nr("16. Bhui K et al. (2003). Ethnic variations in pathways to MHS in UK: SR. Br J "
    "Psychiatry, 182:105–116. doi:10.1192/bjp.182.2.105 — PMID 12562737", size=10)
nr("17. Sekhon M, Cartwright M, Francis JJ (2017). Acceptability of healthcare "
    "interventions: An overview of reviews. BMC Health Serv Res, 17:88. "
    "doi:10.1186/s12913-017-2031-8 — PMID 28126032", size=10)
nr("18. Musiat P, Goldstone P, Tarrier N (2014). Understanding the acceptability of "
    "e-mental health. BMC Psychiatry, 14:109. doi:10.1186/1471-244X-14-109 — "
    "PMID 24725765", size=10)
nr("19. Anderson JK et al. (2017). A scoping literature review of service-level barriers "
    "for access and engagement with MHS for children. Child Youth Serv Rev, 77:164–176. "
    "doi:10.1016/j.childyouth.2017.04.017", size=10)
nr("20. Brown J, Cochrane R, Cardon D (1999). Running large-scale stress workshops for "
    "the general public. J Ment Health, 8:391–402. doi:10.1080/09638239917256", size=10)
nr("21. Brown JS, Elliott SA, Boardman J, Ferns J, Morrison J (2004). Meeting the unmet "
    "need for depression services with psychoeducational self-confidence workshops. "
    "Br J Psychiatry, 185:511–515. doi:10.1192/bjp.185.6.511 — PMID 15572741", size=10)
nr("22. Sclare I, Michelson D, Malpass L, Coster F, Brown J (2015). Innovations in "
    "Practice: DISCOVER CBT workshops for 16-18-year-olds. Child Adolesc Ment Health, "
    "20:102–106. doi:10.1111/camh.12060 — PMID 32680353", size=10)
nr("23. Van Lieshout RJ et al. (2019). Treating Postpartum Depression With 1-Day CBT-"
    "Based Workshops. J Obstet Gynaecol Can, 41:591–592. doi:10.1016/j.jogc.2018.12.004 — "
    "PMID 30922822", size=10)
nr("24. Brown J, Cochrane R, Hancox T (2000). Large scale stress management workshops "
    "for the general public: A controlled evaluation. Behav Cogn Psychother, 28:139–151. "
    "doi:10.1017/S1352465800001053", size=10)
nr("25. Horrell L et al. (2014). One-day CBT self-confidence workshops for people with "
    "depression. Br J Psychiatry, 204:222–233. doi:10.1192/bjp.bp.112.121855 — "
    "PMID 24357574", size=10)
nr("26. Brown JS et al. (2019). School-based early intervention for anxiety and "
    "depression in older adolescents: feasibility RCT of self-referral DISCOVER. J "
    "Adolesc, 71:150–161. doi:10.1016/j.adolescence.2018.11.009 — PMID 30605895", size=10)
nr("27. Van Lieshout RJ et al. (2021). Effect of Online 1-Day CBT-Based Workshops Plus "
    "Usual Care vs Usual Care Alone for Postpartum Depression: RCT. JAMA Psychiatry, "
    "78:1200–1207. doi:10.1001/jamapsychiatry.2021.2488 — PMID 34495280", size=10)
nr("28. Brown JS, Boardman J, Whittinger N, Ashworth M (2010). Can a self-referral "
    "system help improve access to psychological treatments? Br J Gen Pract, 60:365–371. "
    "doi:10.3399/bjgp10X502115 — PMID 20423578", size=10)
nr("29. Sisley E et al. (2011). 'You swim or you sink…I'm still swimming': African "
    "Caribbean women's experiences. Health Soc Care Community, 19:392–402. "
    "doi:10.1111/j.1365-2524.2010.00985.x — PMID 21276117", size=10)
nr("30. Apolinário-Hagen J, Kemper J, Stürmer C (2017). Public Acceptability of E-Mental "
    "Health Treatment Services: A Scoping Review. JMIR Ment Health, 4:e10. "
    "doi:10.2196/mental.6186 — PMID 28373153", size=10)
nr("31. Middendorf J, Kalish A (1996). The 'Change-up' in lectures. Natl Teach Learn "
    "Forum, 5:1–5. doi:10.1002/ntlf.10026", size=10)
nr("32. Bisson JI et al. (2007). Psychological treatments for chronic PTSD: SR & MA. "
    "Br J Psychiatry, 190:97–104. doi:10.1192/bjp.bp.106.021402 — PMID 17267924", size=10)
nr("33. Moreno C et al. (2020). How mental health care should change as a consequence "
    "of the COVID-19 pandemic. Lancet Psychiatry, 7:813–824. "
    "doi:10.1016/S2215-0366(20)30307-2 — PMID 32682460", size=10)
nr("34. Van Dam NT et al. (2013). Establishing a trait anxiety threshold. Anxiety "
    "Stress Coping, 26:70–86. doi:10.1080/10615806.2011.643457 — PMID 22217221", size=10)
nr("35. Brown JS et al. (2010). Are self-referrers just the worried well? Soc "
    "Psychiatry Psychiatr Epidemiol, 45:396–401. doi:10.1007/s00127-009-0079-4 — "
    "PMID 19536444", size=10)
nr("36. Watkins E et al. (2000). Meeting the needs of people with common mental "
    "disorders. J Ment Health, 9:445–456. doi:10.1080/jmh.9.4.445.456", size=10)
nr("37. Woodall A et al. (2010). Barriers to participation in MH research. BMC "
    "Psychiatry, 10:103. doi:10.1186/1471-244X-10-103 — PMID 21126334", size=10)
nr("38. Gronholm PC, Nye E, Michelson D (2018). Stigma related to targeted school-based "
    "MH interventions: SR of qualitative evidence. J Affect Disord, 240:17–26. "
    "doi:10.1016/j.jad.2018.07.023 — PMID 30041074", size=10)
nr("39. Pill R, Prior L, Wood F (2001). Lay attitudes to professional consultations: "
    "Br Med Bull, 57:207–219. doi:10.1093/bmb/57.1.207 — PMID 11719917", size=10)
nr("40. Beck AT, Rush AJ, Shaw BF, Emery G (1979). Cognitive Therapy of Depression. "
    "Guilford Press, NY. ISBN 978-0-89862-919-4 (sách kinh điển CBT)", size=10)
nr("41. Brown JSL (2018). Increasing access to psychological treatments for adults: "
    "Rationale and lessons from the UK. Int J Ment Health Syst, 12:67. "
    "doi:10.1186/s13033-018-0244-9 — PMID 30450123", size=10)
nr("42. O'connor PJ, Martin B, Weeks CS, Ong L (2014). Factors that influence young "
    "people's MH help-seeking: Health Belief Model study. J Adv Nurs, 70:2577–2587. "
    "doi:10.1111/jan.12423 — PMID 24720449", size=10)
nr("43. House J, Marasli P, Lister M, Brown JSL (2017). Male views on help-seeking for "
    "depression: A Q methodology study. Psychol Psychother Theory Res Pr, 91:117–140. "
    "doi:10.1111/papt.12144 — PMID 28452175", size=10)
nr("44. Fennell M (2006). Overcoming Low Self-Esteem: A Self-Help Guide. Robinson, "
    "London. ISBN 978-1-84529-201-6 (sách self-help kinh điển)", size=10)
nr("45. Rickwood D, Deane FP, Wilson CJ, Ciarrochi J (2005). Young people's "
    "help-seeking. Aust E-J Adv Ment Health, 4:218–251. doi:10.5172/jamh.4.3.218", size=10)
nr("46. McKeague L et al. (2018). Exploring feasibility of school-based self-referral "
    "intervention for emotional difficulties in older adolescents. Child Adolesc Ment "
    "Health, 23:198–205. doi:10.1111/camh.12234 — PMID 32677298", size=10)
nr("47. Bowen A et al. (2012). Patterns of Depression and Treatment in Pregnant and "
    "Postpartum Women. Can J Psychiatry, 57:161–167. doi:10.1177/070674371205700305 — "
    "PMID 22398002", size=10)
nr("48. O'Mahen HA, Flynn HA (2008). Preferences and Perceived Barriers to Treatment "
    "for Depression during Perinatal Period. J Womens Health, 17:1301–1309. "
    "doi:10.1089/jwh.2007.0631 — PMID 18816195", size=10)
nr("49. Chabrol H et al. (2004). Acceptability of psychotherapy and antidepressants for "
    "PND. J Reprod Infant Psychol, 22:5–12. doi:10.1080/02646830310001643003", size=10)
nr("50. Arundell LL et al. (2021). Effectiveness of adapted psychological interventions "
    "for ethnic minority groups: SR. Clin Psychol Rev, 88:102063. "
    "doi:10.1016/j.cpr.2021.102063 — PMID 34265674", size=10)
nr("51. Clark DM et al. (2009). Improving access to psychological therapy: Initial "
    "evaluation of two UK demonstration sites (IAPT). Behav Res Ther, 47:910–920. "
    "doi:10.1016/j.brat.2009.07.010 — PMID 19647230", size=10)
nr("52. Brown JSL et al. (2014). How equitable are psychological therapy services in "
    "South East London now? Soc Psychiatry Psychiatr Epidemiol, 49:1893–1902. "
    "doi:10.1007/s00127-014-0900-6 — PMID 24943282", size=10)
nr("53. Kennedy I et al. (2016). A service evaluation of self-referral to military "
    "mental health teams. Occup Med, 66:394–398. doi:10.1093/occmed/kqw032 — "
    "PMID 27121626", size=10)
nr("54. Andersen RM (1995). Revisiting the behavioral model and access to medical care: "
    "Does it matter? J Health Soc Behav, 36:1–10. doi:10.2307/2137284 — PMID 7738325", size=10)
nr("55. Tomczyk S et al. (2020). Ready, Willing and Able? Investigation of TPB in "
    "Help-Seeking. Prev Sci, 21:749–760. doi:10.1007/s11121-020-01099-2 — "
    "PMID 32067149", size=10)
nr("56. National Collaborating Centre for Mental Health (2010). Depression: Treatment "
    "and Management of Depression in Adults (Updated Edition). British Psychological "
    "Society. NICE Clinical Guideline 90 — https://www.nice.org.uk/guidance/cg90", size=10)
nr("57. The British Psychological Society & The Royal College of Psychiatrists (2010). "
    "https://www.bps.org.uk/", size=10)
nr("58. Iorfino F et al. (2021). Right Care, First Time: Theory-Based Automated "
    "Protocol for Clinical Staging. Front Public Health, 9:1230–1239. "
    "doi:10.3389/fpubh.2021.621862 — PMID 34277529", size=10)
nr("59. Sagar-Ouriaghli I et al. (2020). Engaging male students with mental health "
    "support: qualitative focus group study. BMC Public Heal, 20:1159. "
    "doi:10.1186/s12889-020-09269-1 — PMID 32709251", size=10)
nr("60. Aggarwal NK et al. (2016). Clinician descriptions of communication strategies "
    "for racial/ethnic minorities. Patient Educ Couns, 99:198–209. "
    "doi:10.1016/j.pec.2015.09.001 — PMID 26365276", size=10)
nr("61. Barello S, Castiglioni C, Bonanomi A, Graffigna G (2019). The Caregiving Health "
    "Engagement Scale (CHE-s). BMC Public Heal, 19:1562. "
    "doi:10.1186/s12889-019-7855-1 — PMID 31775694", size=10)

H2("4.2. Nguồn xác minh ngoài (em đã đối chiếu)")
nr("• Bài gốc IJERPH: doi:10.3390/ijerph19052831 — em đã đọc full PDF "
   "02_Papers-goc/UK_BESST_PLACES/Brown_2022_PLACES_IJERPH.pdf (884 KB)", size=11)
nr("• PMC NCBI: PMC8909998 — verbatim verification của expansion P-L-A-C-E-S "
   "(Publicity, Lay, Acceptable, Convenient, Effective, Self-referral)", size=11)
nr("• MDPI Open Access: https://www.mdpi.com/1660-4601/19/5/2831", size=11)
nr("• BESST trial gốc test PLACES Self-referral: PMID 38759665, "
   "doi:10.1016/S2215-0366(24)00101-9", size=11)
nr("• Brown 2018 review (ref 41 trong bài này): Int J Ment Health Syst — bối cảnh "
   "PLACES sớm hơn", size=11)
nr("• Editorial Brown & Carter 2025 (B5 corpus dự án) — tổng hợp PLACES + BESST", size=11)

# =====================================================================
# PHẦN 5 — TRUY VẾT NỘI BỘ
# =====================================================================
H1("PHẦN 5 — TRUY VẾT NỘI BỘ DỰ ÁN")

H2("5.1. RAG (rag_db_full)")
nr("• Bài này CHƯA có chunk RAG riêng — corpus index 47 chunks không bao gồm Brown 2022 "
   "PLACES paper", size=11)
nr("• Liên kết qua chunk RAG khác: QT08 (Wen 2020 hỗ trợ trường), VN030 (Happy House), "
   "INSIGHT_05 (yếu tố bảo vệ)", size=11)
nr("• Đề xuất: index Brown 2022 PLACES vào RAG (cần PDF text extraction → chunking)", size=11)

H2("5.2. KG (06_Scripts/kg_data)")
nr("• Bài này CHƯA có node trong nodes.json", size=11)
nr("• Đề xuất bổ sung node: type=model_paper, country=UK+Canada, "
   "intervention=PLACES_workshops, evidence_level=4 (case series)", size=11)

H2("5.3. Glossary nội bộ")
nr("• PLACES expansion đã được đính chính trong 5 file glossary trong phiên 25/04/2026 "
   "(từ \"Psychoeducational Low-intensity Acceptance Coping Strategies\" sai → "
   "\"Publicity, Lay, Acceptable, Convenient, Effective, Self-referral\" đúng)", size=11)
nr("• Files: 06_Scripts/glossary_data/glossary_v3_pedagogical.json (dòng 140); "
   "06_Scripts/glossary_data/glossary_full.json; "
   "06_Scripts/glossary_data/glossary_enhanced.json; "
   "tro-ly-nghien-cuu-tam-ly/web/data/glossary.json; "
   "tro-ly-nghien-cuu-tam-ly-light/web/data/glossary.json", size=11)

H2("5.4. Doc liên quan trong cùng dòng câu hỏi")
nr("• 01_Bao-cao/BESST_PLACES_giai_thich_cho_thay_25042026_v2.docx — Q&A ngắn về 2 "
   "acronym BESST + PLACES", size=11)
nr("• 01_Bao-cao/Bai_dich_phan_bien/B5_Brown_Carter_2025_dich_phan_bien_25042026.docx "
   "— bài 3 (editorial) đã hoàn thành", size=11)
nr("• 01_Bao-cao/Bai_dich_phan_bien/PLACES_Brown_2022_dich_phan_bien_25042026.docx "
   "(DOC NÀY)", size=11)
nr("• [Tiếp theo] 01_Bao-cao/Bai_dich_phan_bien/BESST_Brown_2024_dich_phan_bien_25042026.docx "
   "— bài 1 BESST trial sẽ hoàn thành sau", size=11)

# =====================================================================
H1("KẾT THÚC DOC")
note("Doc này hoàn thành theo workflow chuẩn (memory feedback_research_workflow.md): "
     "dịch song ngữ EN↔VN từng đoạn của bài Brown 2022 PLACES + phản biện chữ đỏ có "
     "dẫn chứng cụ thể từ corpus + so sánh với 4 mô hình help-seeking khác (Andersen, "
     "TPB, HBM, Gask) + bảng adaptation chi tiết cho VN + reference đầy đủ 61 ref bài "
     "gốc + 6 nguồn xác minh ngoài. Đã kiểm 3 vòng: (1) PLACES expansion P-L-A-C-E-S "
     "verbatim với PMC8909998; (2) số liệu 4 case studies (28→120, 41%, 90-95%, "
     "70%, 64.5%) khớp PDF gốc; (3) reference DOI/PMID đầy đủ.")
note("Sản phẩm tiếp theo: doc dịch + phản biện cho bài 1 (BESST Brown 2024 Lancet "
     "Psychiatry RCT). Em làm tuần tự + báo bác sau mỗi doc để kiểm.",
     color=BLUE)

d.save(OUT)
print('Wrote:', OUT)
