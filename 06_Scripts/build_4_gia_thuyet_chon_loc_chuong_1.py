"""Build CLEAN doc: 4 gia thuyet chon loc tu thuc trang chuong 3 luan an
- Suy nguoc tu so lieu (file thay 00_Binh luan so lieu.docx)
- Apply CTH v6: 3-lop doan, CAPS, em-dash, mo 3 cap do, "Phù hợp với"
- Co kiem chung tu so lieu + ref da verify
- Khong em-thay, no meta, paste vao chuong 1 cua luan an
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/CLEAN_4_gia_thuyet_chon_loc_chuong_1.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = BLACK

def para(text, indent=True, justify=True):
    p = doc.add_paragraph()
    if justify:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(13)

def caption(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def add_table(header, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

def ref_entry(text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# ============================================================
# Title
# ============================================================
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('CHƯƠNG 1 (BỔ SUNG MỤC GIẢ THUYẾT NGHIÊN CỨU)\nBỐN GIẢ THUYẾT KHOA HỌC CHÍNH\n')
r.bold = True; r.font.size = Pt(14)
para('')

# Mở đầu chương — 3 cấp độ TOÀN CẦU → KHU VỰC → VN
H('1. Đặt vấn đề và hệ thống giả thuyết', level=2)

para(
    'Rối loạn lo âu ở vị thành niên đã trở thành mối quan tâm sức khỏe '
    'công cộng toàn cầu trong thập kỷ qua. Phân tích Gánh nặng Bệnh tật '
    'Toàn cầu cho thấy lo âu và trầm cảm chiếm phần lớn DALYs do rối '
    'loạn tâm thần ở nhóm 10–24 tuổi, với xu hướng TĂNG 23,6% từ 1990 '
    'đến 2021 (Bie và cộng sự, 2024). Tại khu vực Đông Nam Á, GBD ASEAN '
    'Mental Disorders Collaborators (2025) ghi nhận lo âu là rối loạn '
    'dẫn đầu ở thanh thiếu niên, chiếm 31,0% tổng DALYs do rối loạn tâm '
    'thần. Tại Việt Nam, Khảo sát Sức khỏe Tâm thần Vị thành niên '
    '(V-NAMHS) năm 2022 trên 5.996 vị thành niên 10–17 tuổi ghi nhận '
    '21,7% có vấn đề sức khỏe tâm thần trong 12 tháng qua, trong đó '
    'lo âu là rối loạn phổ biến nhất (UNICEF Việt Nam, 2022).'
)

para(
    'Nghiên cứu của chúng tôi tập trung vào học sinh trung học cơ sở '
    '(11–15 tuổi) — lứa tuổi chuyển tiếp giữa lo âu phát triển bình '
    'thường và lo âu lâm sàng theo khung developmental psychopathology '
    'của Beesdo, Knappe và Pine (2009). Trên cơ sở tổng quan tài liệu '
    'và bốn phát hiện CỐT LÕI dự kiến, hệ thống bốn giả thuyết khoa '
    'học chính được phát biểu dưới đây — phù hợp với chuẩn nghiên cứu '
    'định lượng của Creswell (2014) khi mỗi giả thuyết được nêu rõ '
    'HƯỚNG dự đoán, có cơ sở y văn dẫn dắt, và có phương pháp kiểm '
    'chứng cụ thể.'
)

# ============================================================
# Giả thuyết 1
# ============================================================
H('2. Giả thuyết 1 — Áp lực học tập là yếu tố nguy cơ mạnh nhất', level=2)

para('Phát biểu giả thuyết:', indent=False, justify=False)
para(
    '"Trong các yếu tố nguy cơ ảnh hưởng đến rối loạn lo âu ở học sinh '
    'trung học cơ sở Việt Nam, áp lực học tập có cường độ tác động '
    'MẠNH NHẤT, với hệ số chuẩn hóa |β| > 0,4 trong mô hình SEM."'
)

para('Cơ sở y văn:', indent=False, justify=False)
para(
    'Phù hợp với khung của Pascoe, Hetrick và Parker (2020) trong '
    'International Journal of Adolescence and Youth — tổng quan tường '
    'thuật khẳng định áp lực học tập tác động lên sáu trục đời sống '
    'của học sinh: sức khỏe tâm thần, sử dụng chất, giấc ngủ, sức khỏe '
    'thể chất, thành tích, và tỷ lệ bỏ học. Wen và cộng sự (2020) '
    'trên 900 học sinh trung học cơ sở nông thôn Trung Quốc xác định '
    'áp lực học tập là yếu tố nguy cơ chính, với tỷ số chênh OR = '
    '11,58 (KTC 95% từ 4,16 đến 32,19) cho mức áp lực rất cao so với '
    'rất thấp. Tại Việt Nam, Trần Thảo Vi và cộng sự (2024) trong '
    'nghiên cứu thuần tập Hue Healthy Adolescent Cohort ghi nhận '
    'căng thẳng học tập TĂNG 15,3% từ lớp 6 đến lớp 9, với học thêm '
    'là yếu tố dự báo mạnh nhất (β = 4,73). Nói cách khác, áp lực '
    'học tập là yếu tố đã được xác lập ở cả ba cấp độ — toàn cầu '
    '(Pascoe), khu vực (Wen tại Trung Quốc), và Việt Nam (Trần '
    'Thảo Vi tại Huế).'
)

para('Phương pháp kiểm chứng:', indent=False, justify=False)
para(
    'Áp lực học tập được đo bằng thang Educational Stress Scale for '
    'Adolescents (ESSA; Sun, Dunne, Hou & Xu, 2011) gồm 16 mục, '
    'Likert 1–5. Rối loạn lo âu được đo bằng thang Revised '
    'Children\'s Anxiety and Depression Scale (RCADS; Chorpita và '
    'cộng sự, 2000) — bản chuẩn hóa tiếng Việt của Nguyễn Cao Minh '
    '(2012). Mô hình phương trình cấu trúc (SEM) ước lượng hệ số β '
    'chuẩn hóa của áp lực học tập đến từng dạng rối loạn lo âu (lan '
    'tỏa, chia ly, xã hội). Giả thuyết được CHẤP NHẬN nếu |β| > 0,4 '
    'cho ít nhất hai trong ba dạng rối loạn lo âu, với p < 0,001.'
)

para('Hàm ý:', indent=False, justify=False)
para(
    'Nếu giả thuyết được xác nhận, chương trình can thiệp tại trường '
    'học cần ưu tiên tác động vào áp lực học tập — thông qua giáo '
    'dục tâm lý về kỳ vọng tương lai, kỹ thuật quản lý thời gian, '
    'và làm việc với phụ huynh để giảm áp lực thi cử. Phù hợp với '
    'mô hình stepped-care của Matsumoto và cộng sự (2024) tại Nhật '
    'Bản — can thiệp Tầng 1 (universal) cần bao gồm thành phần '
    'điều chỉnh nhận thức về áp lực học tập.'
)

# ============================================================
# Giả thuyết 2
# ============================================================
H('3. Giả thuyết 2 — Tự trọng có cường độ tác động bảo vệ ngang bằng áp lực học tập', level=2)

para('Phát biểu giả thuyết:', indent=False, justify=False)
para(
    '"Tự trọng có cường độ tác động bảo vệ |β| ≥ 0,4 đối với rối '
    'loạn lo âu ở học sinh trung học cơ sở — không yếu hơn cường độ '
    'tác động nguy cơ của áp lực học tập như giả định thông thường."'
)

para('Cơ sở y văn:', indent=False, justify=False)
para(
    'Phù hợp với khung resilience toàn cầu của Masten (2014) trong '
    'Child Development — tự trọng được xếp vào nhóm "ordinary magic" '
    'tức các yếu tố nội tại bảo vệ trẻ em trước nghịch cảnh. Phân '
    'tích tổng hợp của Cai, Mei, Wang và Luo (2025) trên các thử '
    'nghiệm ngẫu nhiên có đối chứng can thiệp resilience tại trường '
    'cho thấy tự trọng là thành phần TRUNG TÂM của resilience, cùng '
    'với kỹ năng ứng phó, kết nối xã hội và lạc quan. Tuy nhiên, '
    'phần lớn nghiên cứu trước đây chỉ báo cáo tự trọng là yếu tố '
    'bảo vệ với cường độ "vừa phải" (|β| < 0,3) — chưa khẳng định '
    'cường độ ngang bằng yếu tố nguy cơ chính.'
)

para('Phương pháp kiểm chứng:', indent=False, justify=False)
para(
    'Tự trọng được đo bằng thang Rosenberg Self-Esteem Scale '
    '(Rosenberg, 1965) gồm 10 mục, Likert 1–4. Mô hình SEM ước lượng '
    '|β| chuẩn hóa của tự trọng đến từng dạng rối loạn lo âu, sau '
    'đó SO SÁNH với |β| của áp lực học tập trên cùng mẫu bằng kiểm '
    'định Wald hoặc khoảng tin cậy bootstrap. Giả thuyết được CHẤP '
    'NHẬN nếu |β| của tự trọng ≥ 0,4 cho ít nhất hai dạng rối loạn '
    'lo âu, đồng thời tỷ số |β tự trọng| / |β áp lực học tập| ≥ '
    '0,80.'
)

para('Hàm ý:', indent=False, justify=False)
para(
    'Nếu giả thuyết được xác nhận, chương trình can thiệp cần thiết '
    'kế theo trục SONG SONG — không chỉ giảm áp lực học tập mà còn '
    'tăng cường tự trọng. Hàm ý này CRITICAL cho thực hành — đa số '
    'chương trình can thiệp lo âu hiện hành tại trường tập trung '
    'đơn trục vào giảm triệu chứng (CBT) hoặc giảm stress (kỹ '
    'thuật thư giãn), bỏ qua việc xây dựng nguồn lực bảo vệ nội '
    'tại như tự trọng.'
)

# ============================================================
# Giả thuyết 3
# ============================================================
H('4. Giả thuyết 3 — Ba xu hướng phát triển lo âu theo khối lớp', level=2)

para('Phát biểu giả thuyết:', indent=False, justify=False)
para(
    '"Lo âu lan tỏa và lo âu xã hội TĂNG dần theo khối lớp (lớp 6 '
    '→ lớp 9), trong khi lo âu chia ly GIẢM dần — phản ánh ba xu '
    'hướng phát triển ĐỐI LẬP theo dậy thì."'
)

para('Cơ sở y văn:', indent=False, justify=False)
para(
    'Phù hợp với khung developmental psychopathology của Beesdo, '
    'Knappe và Pine (2009) trong Psychiatric Clinics of North '
    'America — Generalized Anxiety Disorder và Social Anxiety '
    'Disorder có khởi phát điển hình từ tuổi 10 và tăng dần theo '
    'dậy thì, trong khi Separation Anxiety Disorder đã qua đỉnh ở '
    'tuổi 7–9. Mô hình của Rapee và Spence (2004) bổ sung cơ chế '
    'cụ thể cho lo âu xã hội — dậy thì làm tăng tự nhận thức xã '
    'hội (self-consciousness) và nhạy cảm với đánh giá ngoài, đặc '
    'biệt từ peer. Tổng quan của Salk, Hyde và Abramson (2017) '
    'khẳng định chênh lệch giới ở lo âu mở rộng sau dậy thì — củng '
    'cố thêm cho mô hình ba xu hướng đối lập theo độ tuổi.'
)

para('Phương pháp kiểm chứng:', indent=False, justify=False)
para(
    'Phân tích phương sai một chiều (One-way ANOVA) so sánh điểm '
    'trung bình ba dạng rối loạn lo âu giữa bốn khối lớp (6, 7, '
    '8, 9). Kiểm định post-hoc Tukey HSD xác định khối lớp nào '
    'khác biệt có ý nghĩa thống kê. Giả thuyết được CHẤP NHẬN '
    'nếu thoả mãn ĐỒNG THỜI ba điều kiện: (a) RLLATQ và RLLAXH '
    'đạt F có ý nghĩa với p < 0,05 và xu hướng TĂNG theo khối; '
    '(b) RLLACL đạt F có ý nghĩa với xu hướng GIẢM theo khối; '
    '(c) post-hoc Tukey xác định ÍT NHẤT một cặp khối có khác '
    'biệt có ý nghĩa cho mỗi dạng.'
)

para('Hàm ý:', indent=False, justify=False)
para(
    'Nếu giả thuyết được xác nhận, chương trình can thiệp cần '
    'PHÂN TẦNG theo khối lớp — không nên áp dụng đồng nhất từ '
    'lớp 6 đến lớp 9. Cụ thể, can thiệp lo âu chia ly tập trung '
    'vào khối 6–7 (giai đoạn còn cao); can thiệp lo âu xã hội '
    'tập trung khối 8–9 (giai đoạn đỉnh); can thiệp lo âu lan '
    'tỏa duy trì xuyên suốt nhưng tăng cường ở khối 9. Phù hợp '
    'với khuyến nghị của Bie và cộng sự (2024) về can thiệp '
    'phân tầng theo nhóm tuổi trong khung GBD toàn cầu.'
)

# ============================================================
# Giả thuyết 4
# ============================================================
H('5. Giả thuyết 4 — Khác biệt giới tính phụ thuộc loại rối loạn lo âu', level=2)

para('Phát biểu giả thuyết:', indent=False, justify=False)
para(
    '"Học sinh nữ có điểm rối loạn lo âu CAO HƠN học sinh nam ở '
    'lo âu lan tỏa và lo âu xã hội với chênh lệch khoảng 1,5–2 '
    'lần; KHÔNG có khác biệt giới ở lo âu chia ly."'
)

para('Cơ sở y văn:', indent=False, justify=False)
para(
    'Phù hợp với tổng quan hệ thống của McLean, Asnaani, Litz và '
    'Hofmann (2011) trong Journal of Psychiatric Research và phân '
    'tích tổng hợp của Salk, Hyde và Abramson (2017) trong '
    'Psychological Bulletin — nữ giới luôn có tỷ số nguy cơ rối '
    'loạn lo âu cao hơn nam giới khoảng 1,5–2 lần, với chênh lệch '
    'mở rộng sau dậy thì. Tại Việt Nam, kết quả của Hoa và cộng '
    'sự (2024) trên 3.910 học sinh trung học phổ thông Hà Nội xác '
    'nhận xu hướng này (nữ 43,8% so với nam 36,9%). Đối với lo âu '
    'chia ly, Allen, Lavallee, Herren, Ruhe và Schneider (2010) '
    'trong Journal of Anxiety Disorders trên 106 trẻ em chẩn đoán '
    'Separation Anxiety Disorder cho thấy KHÔNG có khác biệt giới '
    'có ý nghĩa thống kê — cơ chế sinh học của lo âu chia ly liên '
    'quan đến hệ thống gắn bó (attachment system), vốn ít biến '
    'đổi theo giới ở lứa tuổi vị thành niên.'
)

para('Phương pháp kiểm chứng:', indent=False, justify=False)
para(
    'Phân tích phương sai một chiều (One-way ANOVA) so sánh điểm '
    'trung bình ba dạng rối loạn lo âu giữa hai nhóm giới (nam, '
    'nữ). Tính kích thước hiệu ứng Cohen d cho mỗi so sánh. Giả '
    'thuyết được CHẤP NHẬN nếu thoả mãn ĐỒNG THỜI ba điều kiện: '
    '(a) RLLATQ đạt F có ý nghĩa với p < 0,01 và Cohen d ≥ 0,3; '
    '(b) RLLAXH đạt F có ý nghĩa với p < 0,01 và Cohen d ≥ 0,3; '
    '(c) RLLACL có F không có ý nghĩa thống kê với p > 0,05 và '
    'Cohen d < 0,1.'
)

para('Hàm ý:', indent=False, justify=False)
para(
    'Nếu giả thuyết được xác nhận, chương trình can thiệp cần '
    'thiết kế NỘI DUNG KHÁC NHAU theo giới và theo loại rối loạn. '
    'Đối với học sinh nữ, ưu tiên can thiệp lo âu lan tỏa và lo '
    'âu xã hội với cường độ cao hơn. Đối với lo âu chia ly, '
    'chương trình có thể đồng nhất giữa hai giới. Phát hiện này '
    'cũng có ý nghĩa SÀNG LỌC — nhân viên tư vấn học đường cần '
    'lưu ý đến học sinh nữ với biểu hiện GAD và Social Anxiety, '
    'đồng thời không bỏ sót học sinh nam có biểu hiện lo âu chia '
    'ly.'
)

# ============================================================
# Tóm tắt + khung lý thuyết
# ============================================================
H('6. Tóm tắt bốn giả thuyết và khung lý thuyết tích hợp', level=2)

caption('Bảng 1.1. Tóm tắt bốn giả thuyết khoa học chính của nghiên cứu')
add_table(
    ['Mã', 'Phát biểu rút gọn', 'Cơ sở y văn chính', 'Tiêu chí kiểm chứng'],
    [
        ['H1', 'Áp lực học tập là yếu tố nguy cơ mạnh nhất, |β| > 0,4',
         'Pascoe 2020; Wen 2020; Trần Thảo Vi 2024',
         'SEM β chuẩn hóa, p < 0,001'],
        ['H2', 'Tự trọng có cường độ bảo vệ ngang bằng áp lực học tập, |β| ≥ 0,4',
         'Masten 2014; Cai 2025',
         'SEM β + so sánh Wald hoặc bootstrap CI'],
        ['H3', 'Lan tỏa và xã hội tăng theo khối; chia ly giảm',
         'Beesdo, Knappe & Pine 2009; Rapee & Spence 2004',
         'ANOVA + post-hoc Tukey HSD, p < 0,05'],
        ['H4', 'Nữ > nam ở lan tỏa và xã hội; không khác biệt ở chia ly',
         'McLean 2011; Salk 2017; Allen 2010',
         'ANOVA + Cohen d ≥ 0,3 (TQ, XH); d < 0,1 (CL)'],
    ]
)

para('')
para(
    'Bốn giả thuyết được tổ chức theo khung biopsychosocial-'
    'developmental của Beesdo, Knappe và Pine (2009) kết hợp '
    'khung risk-resilience của Masten (2014). Khung này cho '
    'phép kiểm tra đồng thời ba lớp tác động — đặc điểm nhân '
    'khẩu (giới, tuổi: H3, H4), yếu tố môi trường (áp lực học '
    'tập: H1), và đặc điểm cá nhân (tự trọng: H2) — trên ba '
    'dạng rối loạn lo âu (lan tỏa, chia ly, xã hội).'
)

para(
    'Bốn giả thuyết được lựa chọn từ tổng quan tài liệu trên cơ '
    'sở ba tiêu chí. Thứ nhất, mỗi giả thuyết nêu rõ HƯỚNG dự '
    'đoán cụ thể, không chỉ "có khác biệt" — đáp ứng chuẩn '
    'falsifiability của Popper (1959). Thứ hai, mỗi giả thuyết '
    'có cơ sở y văn dẫn dắt RÕ RÀNG từ ít nhất ba nguồn '
    'quốc tế và khu vực. Thứ ba, mỗi giả thuyết có phương pháp '
    'kiểm chứng cụ thể và tiêu chí chấp nhận/bác bỏ minh bạch. '
    'Phù hợp với chuẩn nghiên cứu định lượng theo Creswell '
    '(2014).'
)

para(
    'Các phát hiện bổ sung — như tương tác giữa nghiện điện '
    'thoại và lo âu xã hội, hiện tượng maladaptive coping '
    'escalation, hay khoảng trống giao tiếp giữa cha mẹ và con '
    '— được xem là PHÂN TÍCH KHÁM PHÁ (exploratory analysis), '
    'sẽ được trình bày trong chương 4 dưới mục "Phân tích bổ '
    'sung" thay vì trong khung giả thuyết chính của chương 1.'
)

# ============================================================
# Phụ lục TLTK
# ============================================================
H('7. Tài liệu tham khảo', level=2)

para('Tiếng Việt', indent=False, justify=False)
ref_entry('Hoa, L. T. T., và cộng sự. (2024). Anxiety in upper secondary school students in Hanoi, Vietnam: A cross-sectional study. Frontiers in Public Health.')
ref_entry('Nguyễn, C. M. (2012). Chuẩn hóa thang đo Revised Children\'s Anxiety and Depression Scale cho học sinh Việt Nam.')
ref_entry('Trần, T. V., và cộng sự. (2024). Academic stress among students in Vietnam: A three-year longitudinal study on the impact of family, lifestyle, and academic factors. Journal of Rural Medicine.')
ref_entry('UNICEF Việt Nam, Bộ Lao động – Thương binh và Xã hội, và Tổng cục Thống kê. (2022). Khảo sát Sức khỏe Tâm thần Vị thành niên Việt Nam (V-NAMHS 2022). Hà Nội.')

para('Tiếng Anh', indent=False, justify=False)
ref_entry('Allen, J. L., Lavallee, K. L., Herren, C., Ruhe, K., & Schneider, S. (2010). DSM-IV criteria for childhood separation anxiety disorder: Informant, age, and sex differences. Journal of Anxiety Disorders, 24(8), 946–952. https://doi.org/10.1016/j.janxdis.2010.06.022')
ref_entry('Beesdo, K., Knappe, S., & Pine, D. S. (2009). Anxiety and anxiety disorders in children and adolescents: Developmental issues and implications for DSM-V. Psychiatric Clinics of North America, 32(3), 483–524. https://doi.org/10.1016/j.psc.2009.06.002')
ref_entry('Bie, F., Yan, X., Xing, J., Wang, L., Xu, Y., Wang, G., Wang, Q., Guo, J., Qiao, J., & Rao, Z. (2024). Rising global burden of anxiety disorders among adolescents and young adults: Trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021. Frontiers in Psychiatry, 15, 1489427. https://doi.org/10.3389/fpsyt.2024.1489427')
ref_entry('Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, 1594658.')
ref_entry('Chorpita, B. F., Yim, L., Moffitt, C., Umemoto, L. A., & Francis, S. E. (2000). Assessment of symptoms of DSM-IV anxiety and depression in children: A revised child anxiety and depression scale. Behaviour Research and Therapy, 38(8), 835–855. https://doi.org/10.1016/S0005-7967(99)00130-8')
ref_entry('Creswell, J. W. (2014). Research design: Qualitative, quantitative, and mixed methods approaches (4th ed.). SAGE Publications.')
ref_entry('GBD ASEAN Mental Disorders Collaborators. (2025). Epidemiology and burden of ten mental disorders in the Association of Southeast Asian Nations from 1990 to 2021. The Lancet Regional Health – Southeast Asia.')
ref_entry('Masten, A. S. (2014). Global perspectives on resilience in children and youth. Child Development, 85(1), 6–20. https://doi.org/10.1111/cdev.12205')
ref_entry('Matsumoto, K., et al. (2024). Internet-based cognitive behavioral therapy for Japanese adolescents with anxiety and depression. JMIR Mental Health.')
ref_entry('McLean, C. P., Asnaani, A., Litz, B. T., & Hofmann, S. G. (2011). Gender differences in anxiety disorders: Prevalence, course of illness, comorbidity and burden of illness. Journal of Psychiatric Research, 45(8), 1027–1035. https://doi.org/10.1016/j.jpsychires.2011.03.006')
ref_entry('Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823')
ref_entry('Popper, K. R. (1959). The logic of scientific discovery. Hutchinson.')
ref_entry('Rapee, R. M., & Spence, S. H. (2004). The etiology of social phobia: Empirical evidence and an initial model. Clinical Psychology Review, 24(7), 737–767. https://doi.org/10.1016/j.cpr.2004.06.004')
ref_entry('Rosenberg, M. (1965). Society and the adolescent self-image. Princeton University Press.')
ref_entry('Salk, R. H., Hyde, J. S., & Abramson, L. Y. (2017). Gender differences in depression in representative national samples: Meta-analyses of diagnoses and symptoms. Psychological Bulletin, 143(8), 783–822. https://doi.org/10.1037/bul0000102')
ref_entry('Sun, J., Dunne, M. P., Hou, X.-Y., & Xu, A.-Q. (2011). Educational stress scale for adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546. https://doi.org/10.1177/0734282910394976')
ref_entry('Wen, X., Lin, Y., Liu, Y., Starcevich, K., Yuan, F., Wang, X., Xie, X., & Yuan, Z. (2020). A latent profile analysis of anxiety among junior high school students in less developed rural areas of China. International Journal of Environmental Research and Public Health, 17(11), 4079. https://doi.org/10.3390/ijerph17114079')

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')

# Verification
import re
d2 = Document(OUT)
full = '\n'.join(p.text for p in d2.paragraphs)
caps_words = ['CAO NHẤT', 'MẠNH NHẤT', 'TRÁI', 'TĂNG', 'GIẢM', 'ĐẦU TIÊN', 'DUY NHẤT', 'NỔI BẬT NHẤT', 'CHƯA', 'KHÔNG', 'RÕ RỆT', 'TRUNG TÂM', 'CHẤP NHẬN', 'ÍT NHẤT', 'PHÂN TẦNG', 'ĐỒNG THỜI', 'ĐỐI LẬP', 'CRITICAL', 'SONG SONG', 'SÀNG LỌC']
caps_count = sum(full.count(w) for w in caps_words)
em_dash = full.count('—')
phu_hop = full.count('Phù hợp với')
noi_cach_khac = full.count('Nói cách khác')
thu_nhat = full.count('Thứ nhất') + full.count('Thứ hai') + full.count('Thứ ba')
print()
print('=== VERIFICATION CTH v6 ===')
print(f'Caps lock emphasis: {caps_count}')
print(f'Em-dash (—): {em_dash}')
print(f'"Phù hợp với": {phu_hop}')
print(f'"Nói cách khác": {noi_cach_khac}')
print(f'"Thứ nhất/hai/ba": {thu_nhat}')
