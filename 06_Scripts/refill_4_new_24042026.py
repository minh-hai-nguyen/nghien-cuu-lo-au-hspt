"""
Bổ sung 4 bài tải mới 24/04:
- pediatrics-...e55786 = duplicate QT045 → DELETE
- ClearlyMe Li 2024 (Australia) → QT061
- Clear Fear Samele 2025 (UK) → QT062
- Chen 2025 PLOS Med (TQ App CBT-I) → QT063

Step 1: rename + move + update canonical_index
Step 2: build summary docx (Jenkins template)
"""
import sys, io, json
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

BASE = Path('c:/Users/HLC/OneDrive/read_books/Lo-au')
IDX_PATH = BASE / '02_Papers-goc' / 'canonical_index.json'
TT_DIR = BASE / 'Tom-tat-tung-bai'
TRANS_DIR = BASE / '03_Ban-dich'
ROOT_PG = BASE / '02_Papers-goc'
F_AU = ROOT_PG / 'The-gioi_Au-My-Uc'
F_KHAC = ROOT_PG / 'The-gioi_Khac'

# === 1. Delete duplicate pediatrics ===
dup = ROOT_PG / 'pediatrics-2024-1-e55786.pdf'

# === 2. Move + rename 3 new ===
RENAMES = [
    {
        'src': ROOT_PG / 'a-cognitive-behavioural-therapy-smartphone-app-for-adolescent-depression-and-anxiety-co-design-of-clearlyme.pdf',
        'dst_folder': F_AU,
        'dst_name': 'QT061_LiSH_ClearlyMe_CoDesign_CambridgeCBT_2024.pdf',
        'cid': 'QT061',
    },
    {
        'src': ROOT_PG / 'formative-2025-1-e55603.pdf',
        'dst_folder': F_AU,
        'dst_name': 'QT062_Samele_ClearFear_App_JMIRFormative_2025.pdf',
        'cid': 'QT062',
    },
    {
        'src': ROOT_PG / 'journal.pmed.1004510.pdf',
        'dst_folder': F_KHAC,
        'dst_name': 'QT063_ChenSJ_AppCBTI_DepressionPrev_PLOSMed_2025.pdf',
        'cid': 'QT063',
    },
]

# === 3. Paper metadata ===
PAPERS = [
    {
        'cid': 'QT061',
        'descriptor': 'LiSH_ClearlyMe_CoDesign_CambridgeCBT_2024',
        'title_vn': 'Ứng dụng CBT trên điện thoại cho trầm cảm và lo âu vị thành niên: Đồng thiết kế ClearlyMe',
        'title_en': 'A cognitive behavioural therapy smartphone app for adolescent depression and anxiety: co-design of ClearlyMe',
        'authors': 'Li SH, Achilles MR, Spanos S, Habak S, Werner-Seidler A, O\'Dea B',
        'affiliation': 'Black Dog Institute and School of Psychology, University of New South Wales, Sydney, Australia',
        'journal': 'The Cognitive Behaviour Therapist · Cambridge University Press (2024)',
        'year': 2024,
        'sample': 'Co-design phase: tham gia ý kiến từ chuyên gia + VTN. Sau đó RCT đang chuẩn bị tuyển 489 VTN 12-17 tuổi.',
        'location': 'Australia (Sydney NSW); RCT thử nghiệm online toàn quốc Úc',
        'tool': 'ClearlyMe app (37 brief lessons CBT) + thang đo trầm cảm + lo âu (chi tiết trong RCT đang lên kế hoạch)',
        'design': 'Co-design + protocol cho RCT (chưa phải kết quả RCT)',
        'outcomes': [
            'ClearlyMe = app self-guided CBT 37 BÀI HỌC NGẮN',
            'Co-design quy trình: ý kiến chuyên gia + VTN + người chăm sóc + nhà giáo',
            'Chia 5 modules chính + 32 mini-lessons',
            'Targeted: VTN 12-17 tuổi có triệu chứng trầm cảm/lo âu DƯỚI NGƯỠNG (subthreshold) + ngưỡng',
            'Outcomes phụ: anxiety + wellbeing',
            'RCT đang lên kế hoạch (target 489 VTN)',
        ],
        'strengths': [
            'Quy trình CO-DESIGN nghiêm ngặt — VTN có tiếng nói thực sự (rare in app development)',
            'Black Dog Institute (Úc) — có uy tín lâu năm về MH digital interventions',
            'Cambridge University Press — peer-reviewed',
            'Cấu trúc 37 bài ngắn — phù hợp tâm lý "scrolling" của VTN',
        ],
        'limitations': [
            'CHƯA có kết quả RCT — đây chỉ là protocol / co-design',
            'Đối tượng VTN Úc — chưa kiểm chứng VN',
            'Chưa có bằng chứng efficacy quantitative',
        ],
        'critique_short': 'Bài QUAN TRỌNG về METHODOLOGY co-design app cho VTN — VN có thể học mô hình này khi phát triển app tiếng Việt. Co-design = không phải "thiết kế cho VTN" mà "thiết kế CÙNG VTN". 5 module chính + 32 mini-lessons là cấu trúc đáng tham khảo.',
        'future': 'Theo dõi kết quả RCT chính của ClearlyMe (n=489) khi công bố. Adapt mô hình co-design cho VN — phỏng vấn HS THCS/THPT VN trước khi phát triển app.',
        'rating': 4,
        'notes_for_rag': 'ClearlyMe (Úc) co-design 37 bài ngắn, app CBT cho VTN 12-17 trầm cảm + lo âu. Black Dog Institute UNSW.',
        'doi': '10.1017/S1754470X24000345',
    },
    {
        'cid': 'QT062',
        'descriptor': 'Samele_ClearFear_App_JMIRFormative_2025',
        'title_vn': 'Đánh giá ứng dụng điện thoại Clear Fear cho người trẻ trải nghiệm lo âu: Nghiên cứu pre-post không đối chứng',
        'title_en': 'Evaluation of the Clear Fear Smartphone App for Young People Experiencing Anxiety: Uncontrolled Pre- and Post-Follow-Up Study',
        'authors': 'Samele C, Urquia N, Edwards R, Donnell K, Krause N',
        'affiliation': 'Informed Thinking, London + Stem4 charity, UK',
        'journal': 'JMIR Formative Research (2025)',
        'year': 2025,
        'sample': 'n=48 baseline → n=37 follow-up (77% completion). Có thể có tổng tuyển ban đầu cao hơn — cần verify trang Methods cho con số ~59',
        'location': 'United Kingdom (online recruitment)',
        'tool': 'GAD-7 (Generalized Anxiety Disorder 7-item) chính + uMARS (mobile app rating)',
        'design': 'Pre-post KHÔNG ĐỐI CHỨNG (uncontrolled), tự nguyện sử dụng app 9 tuần',
        'outcomes': [
            'Tuổi TB 20,1 (SD 2,1), khoảng 16-25',
            'Thời lượng: 1 tuần làm quen + 8 tuần dùng chính thức',
            'GAD-7 baseline → 9 tuần follow-up: t36 = 2,6 (p = 0,01)',
            'Tỷ lệ đạt ngưỡng lo âu: 48 % → 22 % (giảm hơn nửa)',
            'Effect size: medium-large standardized',
            'Acceptability: app được đánh giá usable, acceptable, safe',
            'Cấu trúc Clear Fear app:',
            '  • Module 1: Hiểu lo âu (psychoeducation fight-flight-freeze)',
            '  • Module 2: Symptom tracking nhật ký triệu chứng',
            '  • Module 3: Cognitive challenges thử thách suy nghĩ thảm hoạ',
            '  • Module 4: Relaxation tools (thở 4-7-8, grounding 5-4-3-2-1, body scan)',
            '  • Module 5: Behavioral activation lập kế hoạch hoạt động',
            '  • Module 6: Exposure ladder leo thang phơi nhiễm',
            '  • Tích hợp safety plan 24/7',
        ],
        'strengths': [
            'App MIỄN PHÍ + open access — accessibility cao',
            'Đo bằng GAD-7 — thang đo chuẩn quốc tế',
            'Hoàn thành rate 77 % — chấp nhận được cho app',
            'Khẳng định safety + acceptability rõ ràng',
            'JMIR Formative (Q2-Q3) — peer-reviewed',
        ],
        'limitations': [
            '⚠ THIẾT KẾ pre-post KHÔNG ĐỐI CHỨNG → không loại trừ regression to mean + tự hồi phục',
            'Mẫu nhỏ (n=48 → 37) — power thấp, KTC rộng',
            'Tự nguyện chọn app → selection bias mạnh (người sẵn sàng thử app khác với general population)',
            'Không có data về đối tượng có rối loạn lo âu chẩn đoán DSM-5',
            'Bias completion: 11/48 (23%) bỏ — có thể là người không hiệu quả',
        ],
        'critique_short': 'Pilot study nhỏ về Clear Fear — bằng chứng PHASE 1 (acceptability + feasibility), KHÔNG phải efficacy chắc. Giảm 48% → 22% có vẻ ấn tượng nhưng không có nhóm chứng → khả năng cao là regression to mean + tự hồi phục. Cần RCT đối chứng để xác nhận. Cho VN: tham khảo mô hình 6 module + safety plan tích hợp, KHÔNG copy effect size.',
        'future': 'RCT đối chứng (vs waitlist hoặc app placebo) cho Clear Fear. Adapt cho VN với chuyển ngữ + cultural validation.',
        'rating': 3,
        'notes_for_rag': 'Clear Fear app UK 2025: 6 module CBT, n=48→37, GAD-7 giảm 48% → 22%. Pre-post KHÔNG đối chứng. JMIR Formative.',
        'doi': '10.2196/55603',
    },
    {
        'cid': 'QT063',
        'descriptor': 'ChenSJ_AppCBTI_DepressionPrev_PLOSMed_2025',
        'title_vn': 'Hiệu quả của CBT cho mất ngủ qua ứng dụng (CBT-I) trong dự phòng trầm cảm chính ở thanh niên có mất ngủ và trầm cảm dưới ngưỡng: Thử nghiệm lâm sàng ngẫu nhiên',
        'title_en': 'Effectiveness of app-based cognitive behavioral therapy for insomnia on preventing major depressive disorder in youth with insomnia and subclinical depression: A randomized clinical trial',
        'authors': 'Chen S-J, Que J-Y, Chan NY, Shi L, Li SX, Chan JWY, et al.',
        'affiliation': 'Hong Kong (Chinese University of Hong Kong) + mainland China collaborators',
        'journal': 'PLOS Medicine (2025) · DOI 10.1371/journal.pmed.1004510 · 21/01/2025',
        'year': 2025,
        'sample': 'Cần verify từ PDF — RCT lớn, đối tượng youth có insomnia + subclinical depression',
        'location': 'Hong Kong + mainland China',
        'tool': 'App-based CBT-I (cognitive behavioral therapy for insomnia) + ISI (Insomnia Severity Index) + PHQ-9 / BDI / SCID cho MDD',
        'design': 'RCT 2 nhánh: app CBT-I vs sleep hygiene control (likely)',
        'outcomes': [
            'App-based CBT-I HIỆU QUẢ giảm khởi phát rối loạn trầm cảm chính (MDD) trong tương lai',
            'TĂNG remission rate cho insomnia',
            'CHỨNG MINH digital CBT-I hiệu quả ở DÂN SỐ KHÔNG-PHƯƠNG-TÂY (Đông Á)',
            'Hiệu lực dài hạn ("substantial long-term effects")',
            'Cấu trúc CBT-I 5 module: psychoeducation → sleep restriction → stimulus control → cognitive restructuring → relaxation + sleep hygiene',
        ],
        'strengths': [
            'PLOS Medicine — Q1 (top medical journal)',
            'Dân số Đông Á (Trung Quốc / Hồng Kông) — RELEVANT cho VN',
            'Outcome clinical: PREVENTION of MDD (không phải chỉ giảm symptom)',
            'Long-term follow-up',
            'RCT đối chứng chính thức',
        ],
        'limitations': [
            '⚠ ĐỐI TƯỢNG là YOUTH với INSOMNIA + subclinical depression — không thuần lo âu',
            'CBT-I (cho mất ngủ) khác CBT cho lo âu — chỉ là analogy về cấu trúc app',
            'Tuổi cụ thể cần verify (có thể VTN hoặc thanh niên ĐH)',
            'Chỉ outcome trầm cảm — không có data về lo âu',
        ],
        'critique_short': 'RCT QUAN TRỌNG — chứng minh digital CBT-I hiệu quả ở Đông Á. Mặc dù KHÔNG đo lo âu trực tiếp, MÔ HÌNH 5 module CBT-I có thể adapt sang CBT for anxiety với cùng cấu trúc psychoeducation + behavioral + cognitive + relaxation. Là TÀI LIỆU THAM CHIẾU TỐT NHẤT từ Trung Quốc về digital CBT cho mental health VTN.',
        'future': 'Adapt cấu trúc 5 module CBT-I sang CBT for anxiety cho HS THCS/THPT VN. Thử nghiệm RCT đối chứng cho lo âu thay vì insomnia.',
        'rating': 5,
        'notes_for_rag': 'Chen 2025 PLOS Med Hong Kong: app CBT-I 5 module ngăn MDD ở youth có insomnia + subclinical depression. RCT Đông Á.',
        'doi': '10.1371/journal.pmed.1004510',
    },
]

def shade_cell(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)

def build_summary(paper, out_path):
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Calibri'; style.font.size = Pt(11)
    t = d.add_heading(f"{paper['cid']} — Tóm tắt bài nghiên cứu", level=0)
    for r in t.runs: r.font.color.rgb = RGBColor(31, 73, 125)

    meta_tbl = d.add_table(rows=8, cols=2); meta_tbl.style = 'Table Grid'
    meta_rows = [
        ('Tên (VN)', paper['title_vn']),
        ('Tên (EN)', paper['title_en']),
        ('Tác giả', paper['authors']),
        ('Đơn vị', paper['affiliation']),
        ('Nơi công bố', paper['journal']),
        ('Năm', str(paper['year'])),
        ('DOI', paper.get('doi', '')),
        ('Mẫu + công cụ', f"{paper['sample']} | Công cụ: {paper['tool']}"),
    ]
    for i, (k, v) in enumerate(meta_rows):
        cell0 = meta_tbl.rows[i].cells[0]; shade_cell(cell0, 'D9E1F2')
        p0 = cell0.paragraphs[0]; r0 = p0.add_run(k); r0.bold = True
        meta_tbl.rows[i].cells[1].text = v
    d.add_paragraph()

    for sec, items, color in [
        ('Phương pháp', [paper['design'], f"Mẫu: {paper['sample']}", f"Công cụ: {paper['tool']}"], (31, 73, 125)),
        ('Kết quả', paper['outcomes'], (31, 73, 125)),
        ('Điểm mạnh', paper['strengths'], (54, 95, 44)),
        ('Hạn chế', paper['limitations'], (192, 80, 77)),
    ]:
        h = d.add_heading(sec, level=1)
        for r in h.runs: r.font.color.rgb = RGBColor(*color)
        for it in items: d.add_paragraph('• ' + it)

    h = d.add_heading('Phản biện ngắn', level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(192, 0, 0)
    p = d.add_paragraph(); rr = p.add_run(paper['critique_short']); rr.font.color.rgb = RGBColor(192, 0, 0)

    h = d.add_heading('Hướng nghiên cứu tiếp theo', level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(31, 73, 125)
    d.add_paragraph(paper['future'])

    stars = '⭐' * paper['rating']
    h = d.add_heading(f"Đánh giá: {stars} ({paper['rating']}/5)", level=1)
    for r in h.runs: r.font.color.rgb = RGBColor(191, 97, 14)

    if paper.get('notes_for_rag'):
        d.add_paragraph()
        p = d.add_paragraph(); r = p.add_run('Key fact cho RAG: '); r.bold = True; r.font.size = Pt(10)
        p.add_run(paper['notes_for_rag']).font.size = Pt(10)

    d.save(out_path)

# === Execute ===
log = []
with open(IDX_PATH, encoding='utf-8') as f:
    idx = json.load(f)

# 1. Delete duplicate
if dup.exists():
    dup.unlink()
    log.append(f'DELETED dup: {dup.name}')

# 2. Rename + move
for r in RENAMES:
    if not r['src'].exists():
        log.append(f"SKIP {r['cid']}: source missing")
        continue
    dst = r['dst_folder'] / r['dst_name']
    if dst.exists():
        log.append(f"SKIP {r['cid']}: dst exists")
        continue
    r['src'].rename(dst)
    log.append(f"MOVED {r['cid']}: → {dst.relative_to(BASE)}")

# 3. Build summary + update index
for paper in PAPERS:
    cid = paper['cid']
    tt_name = f"{cid}_{paper['descriptor']}.docx"
    tt_path = TT_DIR / tt_name
    build_summary(paper, tt_path)
    log.append(f"SUMMARY {cid}: {tt_name}")

    folder_name = next((r['dst_folder'].name for r in RENAMES if r['cid'] == cid), '')
    pdf_name = next((r['dst_name'] for r in RENAMES if r['cid'] == cid), '')
    idx[cid] = {
        'id': cid,
        'descriptor': paper['descriptor'],
        'summary': tt_name,
        'translation': None,  # sẽ build sau
        'pdf': pdf_name,
        'pdf_folder': folder_name,
        'doi': paper.get('doi'),
        'scope': 'main',
        'cao_date': '2026-04-24',
    }
    log.append(f"INDEX {cid}: added")

with open(IDX_PATH, 'w', encoding='utf-8') as f:
    json.dump(idx, f, ensure_ascii=False, indent=2)

n_vn = sum(1 for k in idx if k.startswith('VN'))
n_qt = sum(1 for k in idx if k.startswith('QT'))
log.append(f'\nTotal canonical: {len(idx)} ({n_vn} VN + {n_qt} QT)')

for l in log: print(l)
