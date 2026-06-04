"""Build doc tra loi: Cutoff DASS-42 la diem cat phai khong?
Format: CAU TRA LOI to xanh truoc phu luc.
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Cutoff_DASS_42_la_diem_cat.docx')

doc = Document()
for s in doc.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)

style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

def H(text, level=1, color=NAVY):
    p = doc.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:18, 2:14, 3:12}.get(level, 11))
    r.font.color.rgb = color

def para_blue(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLUE
    r.font.size = Pt(12); r.bold = bold

def para_black(text, bold=False, italic=False, size=12):
    p = doc.add_paragraph()
    r = p.add_run(text); r.font.color.rgb = BLACK
    r.font.size = Pt(size); r.bold = bold; r.italic = italic

def bullet_blue(text):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(12)

def bullet_black(text, italic=False, size=12):
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.color.rgb = BLACK; r.font.size = Pt(size); r.italic = italic

def add_table(header, rows):
    tbl = doc.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(11); run.font.name = 'Times New Roman'

# ===========================================================
H('Cutoff DASS-42 là điểm cắt phải không?', level=1)
H('— Xác nhận thuật ngữ + bảng cutoff đầy đủ + áp dụng cho VN006 —', level=2)

# CÂU HỎI
H('Câu hỏi của thầy', level=2)
para_blue('Điểm cutoff DASS-42 là điểm cắt phải không Em?')

# BACKGROUND
H('1. Khái niệm "cutoff" trong tâm lý đo lường', level=2)
para_black(
    'ĐÚNG ạ — "cutoff" trong tâm lý đo lường tương đương ba thuật ngữ thường dùng '
    'trong tiếng Việt: ĐIỂM CẮT, NGƯỠNG, hoặc NGƯỠNG PHÂN LOẠI. Cả ba đều chỉ '
    'GIÁ TRỊ ĐIỂM SỐ trên một thang đo, dùng để PHÂN CHIA mẫu thành các nhóm '
    'theo mức độ trầm trọng (severity).'
)
para_black(
    'Một thang đo có thể có nhiều cutoff — ví dụ DASS-42 có 4 cutoff cho mỗi '
    'subscale, chia thành 5 mức từ "bình thường" đến "rất nặng". Khác với '
    'cutoff CHẨN ĐOÁN (diagnostic threshold, dạng nhị phân "có/không có rối '
    'loạn"), cutoff DASS là cutoff PHÂN TẦNG (categorical threshold) chia mức '
    'độ liên tục thành các nhóm.'
)

H('2. Bảng cutoff DASS-42 đầy đủ (Lovibond & Lovibond, 1995)', level=2)
para_black(
    'Thang DASS-42 (Depression Anxiety Stress Scales — bản đầy đủ 42 mục) do '
    'Lovibond và Lovibond (1995) phát triển có ba subscale, mỗi subscale 14 '
    'mục với điểm Likert 0–3. Tổng điểm tối đa mỗi subscale là 42. Cutoff cho '
    'từng subscale như sau:'
)

para_black('Subscale TRẦM CẢM (Depression):', bold=True)
add_table(
    ['Mức', 'Điểm tổng'],
    [
        ['Bình thường (Normal)', '0 – 9'],
        ['Nhẹ (Mild)', '10 – 13'],
        ['Vừa (Moderate)', '14 – 20'],
        ['Nặng (Severe)', '21 – 27'],
        ['Rất nặng (Extremely severe)', '≥ 28'],
    ]
)
para_black('')

para_black('Subscale LO ÂU (Anxiety):', bold=True)
add_table(
    ['Mức', 'Điểm tổng'],
    [
        ['Bình thường (Normal)', '0 – 7'],
        ['Nhẹ (Mild)', '8 – 9'],
        ['Vừa (Moderate)', '10 – 14'],
        ['Nặng (Severe)', '15 – 19'],
        ['Rất nặng (Extremely severe)', '≥ 20'],
    ]
)
para_black('')

para_black('Subscale CĂNG THẲNG (Stress):', bold=True)
add_table(
    ['Mức', 'Điểm tổng'],
    [
        ['Bình thường (Normal)', '0 – 14'],
        ['Nhẹ (Mild)', '15 – 18'],
        ['Vừa (Moderate)', '19 – 25'],
        ['Nặng (Severe)', '26 – 33'],
        ['Rất nặng (Extremely severe)', '≥ 34'],
    ]
)
para_black('')

H('3. Khác biệt với DASS-21 (bản rút gọn)', level=2)
para_black(
    'DASS-21 là bản rút gọn 21 mục (7 mục mỗi subscale thay vì 14). Để giữ '
    'cùng dải cutoff với DASS-42, điểm raw của DASS-21 phải được NHÂN VỚI 2 '
    'trước khi so cutoff. Ví dụ: học sinh đạt 8 điểm raw lo âu trên DASS-21 '
    '→ ×2 = 16 → đối chiếu cutoff lo âu DASS-42 → mức "nặng" (15–19).'
)
para_black(
    'Đây là điểm dễ nhầm lẫn khi đối chiếu giữa các nghiên cứu — một số bài '
    'báo cáo điểm raw DASS-21 (không nhân 2) và áp dụng cutoff DASS-42 trực '
    'tiếp, dẫn đến phân loại SAI mức độ.'
)

H('4. Áp dụng cho bài VN006 Trần Thị Mỵ Lương và Đặng Đức Anh (2020)', level=2)
para_black(
    'Bài VN006 trong cơ sở dữ liệu dự án sử dụng DASS-42 trên 540 học sinh '
    'một trường THPT Chuyên, với tỷ lệ rối loạn lo âu 14,2%. Tuy nhiên, bài '
    'KHÔNG nêu cutoff cụ thể đã áp dụng — không rõ tác giả dùng cutoff gốc '
    'của Lovibond và Lovibond (1995), cutoff bản Việt hóa, hay cutoff tự '
    'đặt. Hậu quả: tỷ lệ 14,2% không so trực tiếp được với các nghiên cứu '
    'khác cùng dùng DASS-21 hoặc DASS-42 nhưng với cutoff khác.'
)
para_black(
    'Khuyến nghị khi đối chiếu hoặc trích dẫn — luôn ghi rõ phiên bản '
    'cutoff được sử dụng. Ví dụ: "tỷ lệ lo âu (DASS-42, cutoff Lovibond '
    'gốc) là 14,2%" thay vì chỉ "tỷ lệ lo âu là 14,2%".'
)

H('5. Phiên bản DASS bản tiếng Việt', level=2)
para_black(
    'DASS-21 đã được Trần Thị Tú Đoan, Trần Tuấn và Fisher (2013) chuẩn hóa '
    'trên phụ nữ trưởng thành ở vùng nông thôn Bắc Việt Nam. Tuy nhiên, '
    'bản chuẩn hóa này TẬP TRUNG VÀO ĐỘ TIN CẬY VÀ ĐỘ GIÁ TRỊ của thang đo '
    '— chứ KHÔNG đặt ra cutoff khác với cutoff Lovibond gốc. Đối với học '
    'sinh trung học cơ sở Việt Nam, hiện CHƯA có bản chuẩn hóa cutoff '
    'riêng — cần áp dụng cutoff Lovibond gốc và ghi rõ giới hạn này khi '
    'báo cáo.'
)

# CÂU TRẢ LỜI
H('6. CÂU TRẢ LỜI', level=2, color=BLUE)
para_blue('Tóm gọn:', bold=True)
bullet_blue('CÓ — "cutoff" trong DASS-42 (và các thang đo tương tự) chính là ĐIỂM CẮT, hay còn gọi là NGƯỠNG PHÂN LOẠI. Ba thuật ngữ này tương đương trong tiếng Việt.')
bullet_blue('DASS-42 có 4 cutoff cho mỗi subscale, chia thành 5 mức: Bình thường — Nhẹ — Vừa — Nặng — Rất nặng. Cutoff khác nhau giữa ba subscale (trầm cảm, lo âu, stress) do tổng điểm tối đa giống nhau (42) nhưng phân phối triệu chứng khác nhau.')
bullet_blue('Đối với LO ÂU: 0–7 bình thường; 8–9 nhẹ; 10–14 vừa; 15–19 nặng; ≥ 20 rất nặng.')
bullet_blue('DASS-21 dùng cùng dải cutoff trên NHƯNG điểm raw phải nhân 2 trước khi so. Đây là nguyên nhân nhầm lẫn phổ biến giữa các nghiên cứu.')

para_blue('Lưu ý quan trọng cho dự án:', bold=True)
bullet_blue('Bài VN006 Trần Thị Mỵ Lương và Đặng Đức Anh (2020) dùng DASS-42 nhưng KHÔNG nêu cutoff cụ thể — đây là điểm yếu phương pháp. Khi trích, ghi rõ giới hạn này.')
bullet_blue('Bản DASS-21 đã chuẩn hóa cho phụ nữ Bắc Việt Nam (Tran, Tran & Fisher, 2013), nhưng CHƯA có bản chuẩn hóa cutoff riêng cho học sinh trung học cơ sở. Khi nghiên cứu trên đối tượng này, áp dụng cutoff Lovibond gốc và ghi rõ giới hạn.')

para_blue('Khi viết báo cáo CTH:', bold=True)
para_blue(
    '"Mức độ rối loạn lo âu được phân loại theo cutoff Lovibond và Lovibond '
    '(1995) cho thang DASS-42 — Bình thường (0–7), Nhẹ (8–9), Vừa (10–14), '
    'Nặng (15–19), và Rất nặng (≥ 20). Đối với DASS-21, điểm raw được nhân '
    'với 2 trước khi đối chiếu cutoff."'
)

# PHỤ LỤC
H('7. Phụ lục — Tài liệu tham khảo', level=2)
para_black(
    '1. Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression '
    'Anxiety Stress Scales (2nd ed.). Sydney: Psychology Foundation of '
    'Australia. — Sách kinh điển đặt ra cutoff DASS-21 và DASS-42.',
    italic=True, size=11
)
para_black(
    '2. Tran, T. D., Tran, T., & Fisher, J. (2013). Validation of the '
    'depression anxiety stress scales (DASS) 21 as a screening instrument '
    'for depression and anxiety in a rural community-based cohort of '
    'northern Vietnamese women. BMC Psychiatry, 13, 24. '
    'https://doi.org/10.1186/1471-244X-13-24 — Bản chuẩn hóa DASS-21 cho '
    'phụ nữ trưởng thành Bắc Việt Nam.',
    italic=True, size=11
)
para_black(
    '3. Henry, J. D., & Crawford, J. R. (2005). The short-form version of '
    'the Depression Anxiety Stress Scales (DASS-21): Construct validity '
    'and normative data in a large non-clinical sample. British Journal '
    'of Clinical Psychology, 44(2), 227–239. '
    'https://doi.org/10.1348/014466505X29657 — Validate DASS-21 trên '
    'mẫu cộng đồng UK.',
    italic=True, size=11
)
para_black(
    '4. Trần Thị Mỵ Lương và Đặng Đức Anh (2020). Rối loạn lo âu ở học '
    'sinh trung học phổ thông [Chuyên]. Tạp chí Khoa học — Trường Đại học '
    'Thủ đô Hà Nội, Số 40/2020, 122–131. [VN006 trong cơ sở dữ liệu dự án.]',
    italic=True, size=11
)

doc.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {OUT.stat().st_size//1024} KB')
