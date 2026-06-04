# -*- coding: utf-8 -*-
"""Fix 9 expert-review issues in VN002 FULL."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')

d = Document(OUT)

# Fixes = list of (find, replace) — applied to every paragraph
FIXES = [
    # 1. Standardize "sức khỏe" → "sức khoẻ" (VN standard per memory)
    ('sức khỏe', 'sức khoẻ'),

    # 2. Anglicism fix: "mental health" standalone → "SKTT" or keep in parens
    # Only fix standalone, not in titled phrases
    ('mental health is a public health', 'sức khoẻ tâm thần là một vấn đề y tế công cộng'),
    (' mental health problem ', ' vấn đề sức khoẻ tâm thần '),
    (' mental health problems ', ' các vấn đề sức khoẻ tâm thần '),

    # 3. Teacher jargon — add definition on first use
    # Find first "subthreshold" and add parenthetical
    ('subthreshold symptoms', 'subthreshold symptoms (triệu chứng dưới ngưỡng)'),
    ('full threshold symptoms', 'full threshold symptoms (triệu chứng đầy đủ ngưỡng)'),
    # mhGAP
    ('mhGAP', 'mhGAP (Mental Health Gap Action Programme — WHO)'),
    # weighting
    ('weighting', 'weighting (gia trọng thống kê)'),

    # 4. Researcher: "đại diện quốc gia" - add weighting mention (only if not already)
    # Handled in specific paragraph additions below

    # 5. Hedge overly assertive claims
    (' — lại khẳng định khác biệt sàng lọc vs chẩn đoán',
     ' — lại một lần nữa phản ánh khác biệt sàng lọc vs chẩn đoán'),
    ('trích dẫn trong mọi báo cáo Lo âu',
     'trích dẫn trong các báo cáo Lo âu'),
    ('cho mọi đề cương nghiên cứu SKTT VTN VN',
     'cho các đề cương nghiên cứu SKTT VTN VN'),
    ('đã từng cite từ VN002 bị thay đổi',  # other edits
     'đã từng cite từ VN002 thay đổi'),
]

applied = 0
# Apply to paragraphs
for p in d.paragraphs:
    t = p.text
    if not t.strip():
        continue
    new_t = t
    changes_here = False
    for find, rep in FIXES:
        if find in new_t and rep not in new_t.replace(find, ''):  # avoid double-apply
            # For jargon definitions — only add on first occurrence in paragraph
            if '(' in rep and '(' not in find:
                # Only apply once per paragraph
                idx = new_t.find(find)
                if idx >= 0 and rep[:20] not in new_t[:idx]:
                    new_t = new_t[:idx] + rep + new_t[idx+len(find):]
                    changes_here = True
            else:
                if find in new_t:
                    new_t = new_t.replace(find, rep)
                    changes_here = True
    if changes_here and new_t != t:
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = new_t
        applied += 1

# Apply to table cells
for table in d.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                t = p.text
                if not t.strip():
                    continue
                new_t = t
                changes_here = False
                for find, rep in FIXES:
                    if find in new_t:
                        if '(' in rep and '(' not in find:
                            idx = new_t.find(find)
                            if idx >= 0:
                                new_t = new_t[:idx] + rep + new_t[idx+len(find):]
                                changes_here = True
                        else:
                            new_t = new_t.replace(find, rep)
                            changes_here = True
                if changes_here and new_t != t:
                    for run in p.runs:
                        run.text = ''
                    if p.runs:
                        p.runs[0].text = new_t
                    applied += 1

d.save(OUT)
print(f'Applied fixes: {applied} paragraphs modified')

# Now apply "đại diện quốc gia" fix — add weighting mention to first occurrence
d2 = Document(OUT)
for p in d2.paragraphs:
    if 'đại diện quốc gia' in p.text and 'gia trọng' not in p.text and 'weighting' not in p.text.lower():
        # Add note at end
        new = p.text + ' (kết quả được gia trọng theo phân bố tuổi–giới–đô thị/nông thôn)'
        for run in p.runs:
            run.text = ''
        if p.runs:
            p.runs[0].text = new
        print(f'  + Added weighting mention to: "...{p.text[:80]}..."')
        break  # only first
d2.save(OUT)

# Verify
d3 = Document(OUT)
txt = '\n'.join(p.text for p in d3.paragraphs)
for t in d3.tables:
    for r in t.rows:
        for c in r.cells:
            txt += ' ' + c.text

print('\nVerification:')
checks = [
    ('sức khỏe (wrong)', txt.count('sức khỏe')),
    ('sức khoẻ (correct)', txt.count('sức khoẻ')),
    ('subthreshold symptoms (triệu chứng dưới ngưỡng)', txt.count('subthreshold symptoms (triệu chứng dưới ngưỡng)')),
    ('mhGAP (Mental Health Gap', txt.count('mhGAP (Mental Health Gap')),
    ('gia trọng near đại diện quốc gia', ('đại diện quốc gia' in txt and 'gia trọng' in txt)),
    ('khẳng định khác biệt (should be 0 — softened)', txt.count('khẳng định khác biệt')),
]
for k, v in checks:
    print(f'  {k}: {v}')
