# -*- coding: utf-8 -*-
"""Bảng tổng hợp lỗi đã sửa — bôi đỏ chỗ sửa"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3); s.right_margin = Cm(2)

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

# TIÊU ĐỀ
h = doc.add_heading('BÁO CÁO KIỂM TRA 3 LƯỢT — SO KHỚP VỚI PDF GỐC', 1)
for r in h.runs: r.font.name = 'Times New Roman'

p = doc.add_paragraph()
r = p.add_run('Kiểm tra 10 bài tóm tắt (02–11) so với bản gốc tiếng Anh (PDF). ')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r = p.add_run('Các chỗ bôi đỏ là lỗi đã phát hiện và sửa.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED; r.bold = True

doc.add_paragraph()

# BẢNG TỔNG HỢP KẾT QUẢ
h2 = doc.add_heading('1. Bảng tổng hợp kết quả kiểm tra', 2)
for r in h2.runs: r.font.name = 'Times New Roman'

t = doc.add_table(rows=11, cols=4)
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER

headers = ['Bài', 'Số liệu kiểm tra', 'Kết quả', 'Lỗi đã sửa']
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    shade(c, 'D9E2F3')

data = [
    ['02 Saikia 2023', '11 số liệu', '✅ 11/11 PASS', 'Không lỗi'],
    ['03 Mandaknalli 2021', '14 số liệu', '🔴 SỬA 2 CHỖ', '(1) Nhãn mức GAD\n(2) Dữ liệu hút thuốc'],
    ['04 NSCH 2020', 'Dữ liệu web', '✅ PASS', 'Không lỗi'],
    ['05 Alharbi 2019', '13 số liệu', '✅ 13/13 PASS', 'Không lỗi (34% = format khác)'],
    ['06 Nakie 2022', '14 số liệu', '✅ 14/14 PASS', 'Không lỗi'],
    ['07 Chen 2023', '15 số liệu', '✅ 15/15 PASS', 'Không lỗi'],
    ['08 Wen 2020', '10 số liệu', '✅ 10/10 PASS', 'Không lỗi'],
    ['09 Qiu 2022', '9 số liệu', '✅ 9/9 PASS', 'Không lỗi'],
    ['10 Xu 2021', '9 số liệu', '🔴 SỬA 1 CHỖ', 'Câu nhầm "3 NC nam>nữ"'],
    ['11 Bhardwaj 2020', '17 số liệu', '✅ 17/17 PASS', 'Không lỗi'],
]

for ri, rd in enumerate(data):
    for ci, v in enumerate(rd):
        c = t.rows[ri+1].cells[ci]
        c.text = v
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if '🔴' in v or 'SỬA' in v:
                    r.font.color.rgb = RED; r.bold = True

doc.add_paragraph()

# CHI TIẾT LỖI ĐÃ SỬA
h3 = doc.add_heading('2. Chi tiết các lỗi đã sửa (bôi đỏ)', 2)
for r in h3.runs: r.font.name = 'Times New Roman'

# Lỗi 1: Mandaknalli nhãn GAD
p = doc.add_paragraph()
r = p.add_run('LỖI 1 — Bài 03 Mandaknalli: Nhãn mức độ lo âu trong bảng')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

t2 = doc.add_table(rows=5, cols=3)
t2.style = 'Table Grid'
h2data = ['Điểm GAD', 'Script cũ (SAI)', 'PDF gốc (ĐÚNG — đã sửa)']
for i, h in enumerate(h2data):
    c = t2.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    shade(c, 'D9E2F3')

rows2 = [
    ['0–4', 'Tối thiểu (Minimal)', 'Nhẹ (Mild) ← SỬA'],
    ['5–9', 'Nhẹ (Mild)', 'Trung bình (Moderate) ← SỬA'],
    ['10–14', 'Trung bình (Moderate)', 'Nặng (Severe) ← SỬA'],
    ['15–21', 'Nặng (Severe)', 'Rất nặng (Very severe) ← SỬA'],
]
for ri, rd in enumerate(rows2):
    for ci, v in enumerate(rd):
        c = t2.rows[ri+1].cells[ci]; c.text = v
        for p in c.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'; r.font.size = Pt(10)
                if 'SỬA' in v:
                    r.font.color.rgb = RED; r.bold = True

p = doc.add_paragraph()
r = p.add_run('Ghi chú: Bản thân PDF gốc mâu thuẫn nội bộ — Abstract viết "minimal, mild, moderate, severe" (theo chuẩn GAD-7 quốc tế) nhưng Table 1 viết "Mild, Moderate, Severe, Very severe". Đã sửa theo Table 1 của PDF gốc.')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True

doc.add_paragraph()

# Lỗi 2: Mandaknalli hút thuốc
p = doc.add_paragraph()
r = p.add_run('LỖI 2 — Bài 03 Mandaknalli: Dữ liệu hút thuốc đảo ngược')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

t3 = doc.add_table(rows=2, cols=2)
t3.style = 'Table Grid'
for i, h in enumerate(['Script cũ (SAI)', 'PDF gốc (ĐÚNG — đã sửa)']):
    c = t3.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    shade(c, 'D9E2F3')

c = t3.rows[1].cells[0]
c.text = 'Hút thuốc lá: 14,8% trong nhóm lo âu'
for p in c.paragraphs:
    for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

c = t3.rows[1].cells[1]
pp = c.paragraphs[0]
r = pp.add_run('Hút thuốc lá (có): 85,2% trong nhóm lo âu')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.color.rgb = RED; r.bold = True

p = doc.add_paragraph()
r = p.add_run('Ghi chú: Script cũ báo cáo 14,8% — đây là tỷ lệ KHÔNG hút thuốc (16/108). PDF Table 3 ghi Yes=416 (92,44%), anxious smokers=92 (85,19%). Đã sửa thành 85,2% + ghi chú "số liệu 92,4% hút thuốc đáng ngờ".')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True

doc.add_paragraph()

# Lỗi 3: Xu "3 NC"
p = doc.add_paragraph()
r = p.add_run('LỖI 3 — Bài 10 Xu: Câu nhầm "3 nghiên cứu nam > nữ"')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

t4 = doc.add_table(rows=2, cols=2)
t4.style = 'Table Grid'
for i, h in enumerate(['Script cũ (SAI)', 'Đã sửa (ĐÚNG)']):
    c = t4.rows[0].cells[i]; c.text = h
    for p in c.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    shade(c, 'D9E2F3')

c = t4.rows[1].cells[0]
c.text = '"Đây là một trong 3 nghiên cứu (cùng Saikia 2023, Wen 2020) cho thấy nam > nữ"'
for p in c.paragraphs:
    for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

c = t4.rows[1].cells[1]
pp = c.paragraphs[0]
r = pp.add_run('"Trong 11 bài NC, chỉ có 2 nghiên cứu cho thấy nam > nữ: Saikia 2023 và Xu 2021"')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.font.color.rgb = RED; r.bold = True

p = doc.add_paragraph()
r = p.add_run('Ghi chú: Wen 2020 bị nhầm — PDF gốc viết rõ OR nam = 0,262 (nam chỉ bằng 26,2% so nữ), "females are more likely to be moderate or severe anxiety than males". Wen 2020 = NỮ > NAM, không phải nam > nữ.')
r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.italic = True

doc.add_paragraph()

# TỔNG KẾT
h4 = doc.add_heading('3. Tổng kết', 2)
for r in h4.runs: r.font.name = 'Times New Roman'

p = doc.add_paragraph()
r = p.add_run('Tổng số bài kiểm tra: 10\n')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
r = p.add_run('Tổng số liệu so khớp với PDF gốc: 112+\n')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True
r = p.add_run('Bài không lỗi: 8/10 (02, 04, 05, 06, 07, 08, 09, 11)\n')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r = p.add_run('Bài có lỗi đã sửa: 2/10 (03 Mandaknalli — 2 lỗi, 10 Xu — 1 lỗi)\n')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED; r.bold = True
r = p.add_run('Kiểm tra 3 lượt: Round 1 (keyword match), Round 2 (critical numbers), Round 3 (critical claims)\n')
r.font.name = 'Times New Roman'; r.font.size = Pt(12)
r = p.add_run('Tất cả số liệu đã được xác nhận khớp với PDF gốc sau khi sửa.')
r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

doc.save('BANG_LOI_DA_SUA.docx')
sys.stderr.write('BANG_LOI_DA_SUA.docx OK\n')
