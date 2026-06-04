# -*- coding: utf-8 -*-
"""
verify_numbers_v1.py — Quet moi so lieu dinh luong trong bao cao v5-2
va trace nguoc ve ban dich/tom tat tuong ung.

Logic:
1. Parse bao cao -> lay moi paragraph/table chua (OR|AOR|d|g|SMD|SUCRA|%|n=|beta)
2. Voi moi so, tim ten tac gia/canonical ID gan nhat trong paragraph
3. Mo file {ID}_*.docx tuong ung (sum + trans) + grep so
4. Neu khong match -> flag as SUSPECTED FABRICATION
5. Report out
"""
import os, re, sys, io, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
REPORT = os.path.join(ROOT, '01_Bao-cao',
                      'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-4.docx')
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
BD_DIR = os.path.join(ROOT, '03_Ban-dich')

# Map tu ten tac gia/keyword -> canonical ID (de tra cuu ban dich)
AUTHOR_TO_ID = {
    # VN papers
    'Hoa': 'VN001',
    'Hoa 2024': 'VN001',
    'Pham 2024': 'VN003',
    'Ngô Anh Vinh': 'VN015',
    'Vinh 2024': 'VN015',
    'Vinh 2024 Lạng Sơn': 'VN015',
    'Trần Hồ Vĩnh Lộc': 'VN020',
    'Trần Thảo Vi': 'VN021',
    'UNICEF VN 2022': 'VN022',
    'Phạm Thị Ngọc': 'VN025',
    'Đào Thị Ngoãn': 'VN028',
    'Duong 2025': 'VN029',
    'Duong': 'VN029',
    'Duong TPHCM': 'VN029',
    'Tran 2023': 'VN030',
    'Happy House': 'VN030',
    'Tran et al. 2023': 'VN030',
    # QT papers
    'Chen 2023': 'QT007',
    'Wen 2020': 'QT008',
    'Xu 2021': 'QT010',
    'Yin': 'QT013',
    'Zhameden': 'QT013',  # misnamed
    'Zhu 2025': 'QT015',
    'Zhu et al. 2025': 'QT015',
    'Zhu Suzhou': 'QT015',
    'Brunborg': 'QT021',
    'Mojtabai': 'QT023',
    'NHS': 'QT026',
    'UK NHS': 'QT026',
    'Li NMA': 'QT029',
    'Li 2025': 'QT029',
    'Islam': 'QT031',
    'Islam 2025': 'QT031',
    'Cho Korea': 'QT034',
    'Jefferies': 'QT035',
    'Xian': 'QT039',
    'Sasaki': 'QT045',
    'Dong 2025': 'QT047',
    'Dong et al. 2025': 'QT047',
    'Dong PLOS': 'QT047',
    'Hoàng Trung Học': 'VN014',
    'VN014': 'VN014',
    'VN015': 'VN015',
    'Ngô Anh Vinh': 'VN015',
    'VN024': 'VN024',
    'VN026': 'VN026',
    'VN003': 'VN003',
    'QT036': 'QT036',
    'Moon': 'QT036',
    'QT030': 'QT030',
    'Cho': 'QT034',
    'QT034': 'QT034',
    'QT021': 'QT021',
    'QT023': 'QT023',
    'QT010': 'QT010',
    'Stockings': None,  # not in our system
    'Bowlby': None,
    'Festinger': None,
    'Sallis': None,
    'Leventhal': None,
    'Khantzian': None,
    'Dumas': None,
    'JAACAP 2025': None,  # umbrella review, no canonical ID
    'Brown & Carter': None,
}

# Regex cho cac dang so lieu
NUMBER_PATTERNS = [
    (r'OR\s*=\s*(\d+[,.]?\d*)', 'OR'),
    (r'aOR\s*=\s*(\d+[,.]?\d*)', 'aOR'),
    (r'AOR\s*=\s*(\d+[,.]?\d*)', 'AOR'),
    (r'β\s*=\s*[-−]?(\d+[,.]?\d*)', 'beta'),
    (r'd\s*=\s*[-−]?(\d+[,.]?\d*)', 'd'),
    (r'g\s*=\s*[-−]?(\d+[,.]?\d*)', 'g'),
    (r'SMD\s*=\s*[-−]?(\d+[,.]?\d*)', 'SMD'),
    (r'SUCRA\s*=\s*(\d+[,.]?\d*)', 'SUCRA'),
    (r'n\s*=\s*([\d.]+)', 'n'),
    (r'(\d+[,.]\d+)\s*%', 'pct'),
]

def extract_numbers(text):
    """Return list of (metric, value) found in text."""
    out = []
    for pat, metric in NUMBER_PATTERNS:
        for m in re.finditer(pat, text):
            out.append((metric, m.group(1)))
    return out

def normalize(s):
    return s.replace(',', '.').strip()

def find_author_in_text(text):
    """Return list of author keys found in text, sorted by length desc."""
    found = []
    for k in AUTHOR_TO_ID:
        if k in text:
            found.append(k)
    return sorted(found, key=len, reverse=True)

def load_source_text(cid):
    """Load sum + trans text for a canonical ID."""
    text = ''
    for d, label in [(TT_DIR, 'sum'), (BD_DIR, 'trans')]:
        for f in os.listdir(d):
            if f.startswith(cid + '_') and f.endswith('.docx'):
                try:
                    doc = Document(os.path.join(d, f))
                    for p in doc.paragraphs:
                        text += p.text + '\n'
                    for t in doc.tables:
                        for r in t.rows:
                            for c in r.cells:
                                text += c.text + '\n'
                except Exception as e:
                    print(f'WARN cannot read {f}: {e}')
                break
    return text

def number_in_source(val, source_text):
    """Check if value appears in source text, allowing slight format variants."""
    if source_text == '':
        return None  # no source available
    v = normalize(val)
    variants = [val, val.replace(',', '.'), val.replace('.', ',')]
    # Also try removing trailing zeros
    try:
        f = float(v)
        variants.append(f'{f:.2f}'.replace('.', ','))
        variants.append(f'{f:.2f}'.replace('.', '.'))
        variants.append(f'{f:.1f}'.replace('.', ','))
        variants.append(str(int(f)) if f == int(f) else v)
    except ValueError:
        pass
    for vv in variants:
        if vv in source_text:
            return True
    return False

# ============================================================
print(f'Reading report: {os.path.basename(REPORT)}')
doc = Document(REPORT)
issues = []

# Collect all paragraphs + table cells as units
units = []
for i, p in enumerate(doc.paragraphs):
    if p.text.strip():
        units.append(('p', i, p.text))
for ti, t in enumerate(doc.tables):
    for ri, r in enumerate(t.rows):
        row_text = ' | '.join(c.text.strip() for c in r.cells)
        if row_text.strip():
            units.append((f'T{ti}', ri, row_text))

print(f'Total units: {len(units)}')

# Cache source text per ID
source_cache = {}
def get_source(cid):
    if cid not in source_cache:
        source_cache[cid] = load_source_text(cid) if cid else ''
    return source_cache[cid]

# Scan units for numbers
claims_checked = 0
fabricated = []
verified = []
no_source = []
no_author = []

for kind, idx, text in units:
    nums = extract_numbers(text)
    if not nums:
        continue
    authors = find_author_in_text(text)
    if not authors:
        no_author.append((kind, idx, text[:120], nums))
        continue
    # Use nearest/longest author match
    for author_key in authors[:1]:  # only primary
        cid = AUTHOR_TO_ID.get(author_key)
        if cid is None:
            no_source.append((kind, idx, author_key, text[:120], nums))
            continue
        src = get_source(cid)
        if not src:
            no_source.append((kind, idx, author_key, text[:120], nums))
            continue
        for metric, val in nums:
            claims_checked += 1
            found = number_in_source(val, src)
            if found is True:
                verified.append((kind, idx, cid, metric, val))
            elif found is False:
                fabricated.append((kind, idx, cid, author_key, metric, val, text[:160]))

# Report
print('\n' + '='*70)
print('VERIFY REPORT')
print('='*70)
print(f'Total claims with numbers checked: {claims_checked}')
print(f'  VERIFIED in source: {len(verified)}')
print(f'  SUSPECTED FABRICATION: {len(fabricated)}')
print(f'  NO SOURCE FILE (external/MA/non-canonical): {len(no_source)}')
print(f'  NO AUTHOR FOUND (orphan number): {len(no_author)}')

if fabricated:
    print('\n' + '-'*70)
    print('SUSPECTED FABRICATION (needs manual check):')
    print('-'*70)
    for kind, idx, cid, author, metric, val, ctx in fabricated:
        print(f'\n[{kind}:{idx}] {cid} ({author}) — {metric} = {val}')
        print(f'  Context: {ctx}...')

if no_author:
    print('\n' + '-'*70)
    print(f'NO AUTHOR FOUND ({len(no_author)} units) — first 10:')
    print('-'*70)
    for kind, idx, preview, nums in no_author[:10]:
        nums_str = ', '.join(f'{m}={v}' for m, v in nums)
        print(f'[{kind}:{idx}] {preview}')
        print(f'  Numbers: {nums_str}')

# Save JSON report
out_json = os.path.join(os.path.dirname(__file__), 'verify_report_v5-2.json')
with open(out_json, 'w', encoding='utf-8') as f:
    json.dump({
        'report_file': os.path.basename(REPORT),
        'total_checked': claims_checked,
        'verified': len(verified),
        'fabricated': [
            {'kind': k, 'idx': i, 'cid': c, 'author': a, 'metric': m,
             'value': v, 'context': ctx}
            for k, i, c, a, m, v, ctx in fabricated
        ],
        'no_source_count': len(no_source),
        'no_author_count': len(no_author),
    }, f, ensure_ascii=False, indent=2)
print(f'\nReport saved to: {out_json}')
