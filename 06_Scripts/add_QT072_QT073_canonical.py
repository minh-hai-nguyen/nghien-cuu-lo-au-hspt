"""Build 2 Tom-tat + add QT072, QT073 vao canonical + rename PDF + update KG/RAG."""
import sys, io, json, shutil
from pathlib import Path
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

TOMTAT_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/Tom-tat-tung-bai')
PAPERS_DIR = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/02_Papers-goc')
CANONICAL = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/02_Papers-goc/canonical_index.json')
LIGHT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/tro-ly-nghien-cuu-tam-ly-light/web/data')
HEAVY = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/tro-ly-nghien-cuu-tam-ly/web/data')


def build_tomtat(filename, title, paragraphs):
    out = TOMTAT_DIR / filename
    d = Document()
    style = d.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(title); r.bold = True; r.font.size = Pt(14)
    for label, text in paragraphs:
        p = d.add_paragraph()
        if label:
            r = p.add_run(label); r.bold = True; r.font.size = Pt(12); r.add_break()
        r = p.add_run(text); r.font.size = Pt(12)
    d.save(out)
    return out


PAPERS = {
    'QT072': {
        'descriptor': 'Lee_2025_CyberbullyingMeta_TVA',
        'filename': 'QT072_Lee_2025_CyberbullyingMeta_TVA.docx',
        'pdf_old': 'lee-et-al-2025-cyberbullying-victimization-and-mental-health-symptoms-among-children-and-adolescents-a-meta-analysis-of.pdf',
        'pdf_new': 'QT072_Lee_2025_CyberbullyingMeta_TVA.pdf',
        'title': 'QT072 — Tóm tắt: Lee và cộng sự (2025) Meta-analysis dọc cyberbullying',
        'tomtat': [
            ('Tên công trình:',
             '"Cyberbullying Victimization and Mental Health Symptoms Among Children and Adolescents: A Meta-Analysis of Longitudinal Studies" — Jungup Lee, Hoi Shan Cheung, Qiyang Zhang, Rebecca P. Ang (2025), Trauma, Violence, & Abuse, 27(2), 391–406. DOI: 10.1177/15248380241313051. Tạp chí Q1 SAGE.'),
            ('Phương pháp:',
             'META-ANALYSIS DỌC ĐẦU TIÊN về cyberbullying → mental health (các meta-analysis trước chỉ cắt ngang). 27 nghiên cứu thuần tập, N = 27.133 trẻ em + vị thành niên qua 9 quốc gia: Mỹ (12), Trung Quốc (7), Tây Ban Nha (2), Thụy Điển, Hà Lan, Bỉ, Thụy Sĩ, Úc, New Zealand. Random-effects model. Singapore Ministry of Education tài trợ.'),
            ('Kết quả CỐT LÕI:',
             'Cyberbullying → Mental health overall: r = 0,232 (KTC 95% 0,190–0,274); k=27. → Depression r = 0,269 (k=23, KTC 0,211–0,325). → ANXIETY r = 0,229 (k=16, KTC 0,162–0,294). → Loneliness r = 0,170. Tất cả heterogeneity I² = 92,25% (rất cao).'),
            ('Moderator analyses:',
             'Tuổi (β = 0,037, p < 0,001) — tuổi càng cao tác động càng mạnh; Tỷ lệ nữ (β = −0,003, p = 0,006) — nữ % cao hơn → tác động giảm (NGƯỢC giả thuyết); Năm công bố (β = 0,009, p < 0,001) — bài mới hơn có hiệu ứng lớn hơn (cyberbullying ngày càng nghiêm trọng); Văn hóa Tây vs Á: KHÔNG khác biệt; Time interval: KHÔNG khác biệt.'),
            ('Phát hiện ĐẶC BIỆT:',
             'Cyberbullying là vấn đề TOÀN CẦU vượt văn hóa (Western vs Asian KHÔNG khác biệt). Internet "san phẳng" khác biệt văn hóa. Hiệu ứng MẠNH HƠN ở tuổi cao hơn (đỉnh 14-15 tuổi theo Pichel 2021). Hiệu ứng tăng theo năm — phản ánh cyberbullying ngày càng phổ biến.'),
            ('Đánh giá chất lượng:',
             '⭐⭐⭐⭐⭐ Rất cao. Meta-analysis DỌC đầu tiên (chỉ longitudinal, KHÔNG cắt ngang). Trauma, Violence, & Abuse Q1 SAGE. PRISMA. N=27.133 lớn. Đáng tin cậy cao.'),
            ('Ý nghĩa cho VN:',
             'Bổ sung cho Moore 2017 (QT070) bullying truyền thống. Cyberbullying r = 0,229 cho anxiety NHẸ HƠN traditional bullying (Moore OR=1,77). Tham chiếu cho phản biện về NĐT + bắt nạt online — đặc biệt khi luận án CTH chỉ đo "bắt nạt thể chất" (Bảng 3.21), CHƯA đo cyberbullying riêng.'),
        ],
    },
    'QT073': {
        'descriptor': 'He_2025_PowerUpCBTD_China_JAD',
        'filename': 'QT073_He_2025_PowerUpCBTD_China_JAD.docx',
        'pdf_old': '1-s2.0-S0165032725020014-main.pdf',
        'pdf_new': 'QT073_He_2025_PowerUpCBTD_China_JAD.pdf',
        'title': 'QT073 — Tóm tắt: He và cộng sự (2025) Power Up-CBTD pilot RCT China',
        'tomtat': [
            ('Tên công trình:',
             '"Preventing depression in Chinese children and adolescents: A pilot study of a brief school-based cognitive behavioral group program" — Qianyun He, Jina Li, Junhe Wang, Zhiyong Qu (2025), Beijing Normal University. Journal of Affective Disorders. ScienceDirect S0165032725020014. Q1 Elsevier.'),
            ('Phương pháp:',
             'PILOT RCT có đối chứng. Mẫu: 87 HS (lớp 5, 6, 9; tuổi 10-15) screened positive cho trầm cảm (CESD-R ≥ 16). Phân ngẫu nhiên Power Up-CBTD vs treatment-as-usual (TAU). Loại trừ: chậm phát triển nặng, nguy cơ tự sát, đang điều trị MH ngoài trường.'),
            ('Can thiệp Power Up-CBTD:',
             'BRIEF MANUALIZED Cognitive Behavioral Group Therapy (CBGT) tailored cho HS Trung Quốc. Nội dung CBT: tâm lý giáo dục + behavioral activation + tái cấu trúc nhận thức + giải quyết vấn đề + thư giãn + kỹ năng giao tiếp xã hội (American Psychological Association 2019; Garber 2016). Format nhóm — phù hợp môi trường trường học.'),
            ('Công cụ đánh giá:',
             'Trầm cảm: CESD-R (Center for Epidemiologic Studies Depression-Revised) + CDI-S (Children\'s Depression Inventory Short Form). Lo âu: SCARED (Screen for Child Anxiety Related Emotional Disorders). Đo baseline (T0) và post-intervention (T1, trong vòng 2 tuần sau buổi cuối).'),
            ('Kết quả CHÍNH:',
             'Feasibility: attrition 2,3% (1/43); attendance 87,7%; satisfaction M=28,4/32 (SD=4,06). HIỆU QUẢ: trầm cảm CESD-R giảm có ý nghĩa thống kê — adjusted mean difference (AMD) = −10,89 (p < 0,001); CDI-S AMD = −3,23 (p < 0,001). Lo âu (SCARED): chưa rõ trong abstract trang 1, có thể giảm ở mức yếu hơn (đây là chương trình ĐẶC THÙ cho trầm cảm).'),
            ('Đánh giá chất lượng:',
             '⭐⭐⭐⭐ Cao. Pilot RCT, phân ngẫu nhiên, 2 thang đo trầm cảm + 1 thang lo âu, attrition rất thấp 2,3%. Tài trợ: chính sách Healthy China 2030 + Children Adolescents Mental Health Action Plan (2019-2022). Hạn chế: pilot nhỏ (n=87), chưa follow-up dài hạn, hiệu quả lo âu chưa rõ.'),
            ('Ý nghĩa cho VN:',
             'Mô hình tham chiếu CHO TRUNG QUỐC HIỆN ĐẠI (2025). Phù hợp khung tập huấn của thầy: brief CBT group therapy school-based, tailored văn hóa Á. Khả năng adapt cho VN cao do có chung bối cảnh áp lực học tập + gia đình. Cần đọc full text để biết số buổi cụ thể + nội dung từng buổi.'),
        ],
    },
}


def main():
    print('=== Build Tom-tat ===')
    for qt_id, data in PAPERS.items():
        out = build_tomtat(data['filename'], data['title'], data['tomtat'])
        print(f'  ✓ {out.name}')

    print('\n=== Đổi tên PDF ===')
    for qt_id, data in PAPERS.items():
        old = PAPERS_DIR / data['pdf_old']
        new = PAPERS_DIR / data['pdf_new']
        if old.exists() and not new.exists():
            shutil.move(str(old), str(new))
            print(f'  ✓ {data["pdf_old"][:40]}... → {data["pdf_new"]}')
        elif new.exists():
            print(f'  ! {data["pdf_new"]} đã tồn tại — skip')

    print('\n=== Update canonical_index ===')
    with open(CANONICAL, encoding='utf-8') as f:
        can = json.load(f)
    for qt_id, data in PAPERS.items():
        can[qt_id] = {
            'descriptor': data['descriptor'],
            'summary': data['filename'],
            'pdf': data['pdf_new'],
            'pdf_folder': '02_Papers-goc',
        }
    with open(CANONICAL, 'w', encoding='utf-8') as f:
        json.dump(can, f, ensure_ascii=False, indent=2)
    print(f'  +2 entries; total {len(can)} entries')

    print('\n=== Update KG ===')
    new_papers = [
        {'id': 'PAPER_QT072_Lee_2025_CyberbullyingMeta', 'type': 'Paper',
         'label': 'QT072 Lee (2025) Cyberbullying meta-analysis longitudinal r=0,229 anxiety'},
        {'id': 'PAPER_QT073_He_2025_PowerUpCBTD', 'type': 'Paper',
         'label': 'QT073 He (2025) Power Up-CBTD pilot RCT China brief CBGT'},
    ]
    new_concepts = [
        {'id': 'CONCEPT_CYBERBULLYING', 'type': 'Concept',
         'label': 'Cyberbullying victimization (online bullying)'},
        {'id': 'CONCEPT_BRIEF_CBGT', 'type': 'Concept',
         'label': 'Brief CBGT (Cognitive Behavioral Group Therapy)'},
        {'id': 'CONCEPT_CESD_R', 'type': 'Concept',
         'label': 'CESD-R (Center for Epidemiologic Studies Depression-Revised)'},
        {'id': 'CONCEPT_SCARED', 'type': 'Concept',
         'label': 'SCARED (Screen for Child Anxiety Related Emotional Disorders)'},
    ]
    new_edges = [
        ('PAPER_QT072_Lee_2025_CyberbullyingMeta', 'CONCEPT_CYBERBULLYING', 'RELATED_TO'),
        ('PAPER_QT072_Lee_2025_CyberbullyingMeta', 'CONCEPT_SCHOOL_BULLYING', 'RELATED_TO'),
        ('PAPER_QT072_Lee_2025_CyberbullyingMeta', 'CONCEPT_SOCIAL_MEDIA_ADDICTION', 'RELATED_TO'),
        ('PAPER_QT073_He_2025_PowerUpCBTD', 'CONCEPT_BRIEF_CBGT', 'RELATED_TO'),
        ('PAPER_QT073_He_2025_PowerUpCBTD', 'CONCEPT_CESD_R', 'RELATED_TO'),
        ('PAPER_QT073_He_2025_PowerUpCBTD', 'CONCEPT_SCARED', 'RELATED_TO'),
        ('PAPER_QT073_He_2025_PowerUpCBTD', 'TOPIC_INTERVENTION_DESIGN', 'BELONGS_TO'),
    ]
    for label, dir_ in [('LIGHT', LIGHT), ('HEAVY', HEAVY)]:
        path = dir_ / 'questions_kg.json'
        with open(path, encoding='utf-8') as f: kg = json.load(f)
        eids = {n['id'] for n in kg['nodes']}
        eedges = {(e['source'],e['target'],e['relation']) for e in kg['edges']}
        na = 0
        for n in new_papers + new_concepts:
            if n['id'] not in eids: kg['nodes'].append(n); eids.add(n['id']); na += 1
        ea = 0
        for src,tgt,rel in new_edges:
            if (src,tgt,rel) in eedges or src not in eids or tgt not in eids: continue
            kg['edges'].append({'source':src,'target':tgt,'relation':rel}); eedges.add((src,tgt,rel)); ea += 1
        from collections import Counter
        kg['meta']['n_nodes']=len(kg['nodes']); kg['meta']['n_edges']=len(kg['edges'])
        kg['meta']['node_types']=dict(Counter(n['type'] for n in kg['nodes']))
        kg['meta']['edge_relations']=dict(Counter(e['relation'] for e in kg['edges']))
        kg['meta']['updated']='2026-05-10'
        with open(path,'w',encoding='utf-8') as f: json.dump(kg,f,ensure_ascii=False,indent=2)
        nn = len(kg['nodes']); ne = len(kg['edges'])
        print(f'  {label}: +{na} nodes, +{ea} edges; total {nn}/{ne}')

    print('\nDone.')


main()
