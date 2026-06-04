# -*- coding: utf-8 -*-
"""Tóm tắt 2 trang cho 11 bài trong Originals/Translations"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(10)
style.paragraph_format.space_after = Pt(3)
style.paragraph_format.line_spacing = 1.15
for s in doc.sections:
    s.top_margin = Cm(1.5)
    s.bottom_margin = Cm(1.5)
    s.left_margin = Cm(2)
    s.right_margin = Cm(1.5)

def shade(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW')
    we.set(qn('w:w'), str(int(w * 567)))
    we.set(qn('w:type'), 'dxa')
    tcW.append(we)

# TITLE
title = doc.add_heading('', level=0)
r = title.add_run('TÓM TẮT 11 BÀI NGHIÊN CỨU\nLO ÂU VÀ TRẦM CẢM Ở HỌC SINH')
r.font.name = 'Times New Roman'
r.font.size = Pt(14)
r.font.color.rgb = RGBColor(0, 51, 102)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p = doc.add_paragraph()
r = p.add_run('Việt Nam \u2014 Ấn Độ \u2014 Trung Quốc \u2014 Hoa Kỳ \u2014 Ả Rập Saudi \u2014 Ethiopia | 2019\u20132023')
r.font.name = 'Times New Roman'
r.font.size = Pt(9)
r.italic = True
p.alignment = WD_ALIGN_PARAGRAPH.CENTER

# TABLE
data = [
    ['1', 'Jenkins và cs.\n(2023)', 'Mỹ\n(San Diego)', '75', 'PHQ-9A\nGAD-10', 'Lo âu: 50,6%\nTrầm cảm: 44%', 'HS THCS đa sắc tộc. Phương pháp hỗn hợp. Nữ > nam (p=0,002). Phỏng vấn dân tộc học phát hiện bạo lực giới, COVID-19, phân biệt chủng tộc.', '⭐⭐⭐'],
    ['2', 'Saikia và cs.\n(2023)', 'Ấn Độ\n(Assam)', '360', 'DASS-21', 'Lo âu: 24,4%\nTrầm cảm: 22,2%', 'Nam > nữ về lo âu (30% vs 18,9%) — trái y văn. Cha mẹ đơn thân, rượu, game, học kém là yếu tố nguy cơ. Đầu tiên ở ĐB Ấn Độ.', '⭐⭐⭐'],
    ['3', 'Mandaknalli\n& Malusare\n(2021)', 'Ấn Độ', '450', 'Bảng hỏi\nlo âu', 'Nhẹ: 49,4%\nTB: 43,3%\nNặng: 7,3%', 'Nữ lo âu nặng hơn (10,9% vs 3,8%). Thiếu vận động, giấc ngủ kém, hút thuốc là yếu tố nguy cơ.', '⭐⭐'],
    ['4', 'NSCH\n(2020)', 'Hoa Kỳ\n(quốc gia)', '55.162', 'Khảo sát\nquốc gia', 'Lo âu: 16,1%\nTrầm cảm: 8,4%\n(chẩn đoán)', 'Lo âu tăng 61% (2016-2023). Nữ 20,1% vs nam 12,3%. Thiếu điều trị: 61% gặp khó khăn. HS bị ảnh hưởng khó kết bạn gấp 10 lần.', '⭐⭐⭐⭐⭐'],
    ['5', 'Alharbi và cs.\n(2019)', 'Ả Rập\nSaudi', '1.245', 'PHQ-9\nGAD-7', 'Lo âu: 63,5%\nTrầm cảm: 74%', 'Tỷ lệ cực cao (ngưỡng thấp). Nữ > nam (P<0,001). Bài đầu tiên ở Ả Rập Saudi dùng PHQ-9/GAD-7.', '⭐⭐⭐⭐'],
    ['6', 'Nakie và cs.\n(2022)', 'Ethiopia', '849', 'DASS-21', 'Lo âu: 66,7%\nTrầm cảm: 41,4%\nStress: 52,2%', 'Nhai lá khat (AOR=5,6), hút thuốc (AOR=4,8), bệnh mãn tính (AOR=2,1). Đầu tiên ở châu Phi đánh giá cả 3.', '⭐⭐⭐⭐'],
    ['7', 'Chen và cs.\n(2023)', 'Trung Quốc\n(miền Tây)', '63.205', 'PHQ-9\nGAD-7', 'Lo âu: 13,9%\nTrầm cảm: 23%', 'Cỡ mẫu rất lớn. Bắt nạt, giấc ngủ kém, nghiện game là yếu tố nguy cơ. NC lớn đầu tiên miền Tây TQ.', '⭐⭐⭐⭐⭐'],
    ['8', 'Wen và cs.\n(2020)', 'Trung Quốc\n(nông thôn)', '900', 'Thang lo âu\n+ LPA', 'Lo âu nặng:\n24,78%', 'LPA xác định 3 nhóm lo âu. Nam > nữ. Hỗ trợ tâm thần tại trường là yếu tố bảo vệ. Đầu tiên dùng LPA nông thôn TQ.', '⭐⭐⭐⭐'],
    ['9', 'Qiu và cs.\n(2022)', 'Trung Quốc', '2.079', 'SAS\nCES-D\nCD-RISC', 'Lo âu: 13,4%\nTrầm cảm: 26%', 'Nuôi dạy tiêu cực tăng nguy cơ (OR=1,82-2,01). Phục hồi thấp: OR trầm cảm=6,74. Tích cực giảm (OR=0,30-0,32).', '⭐⭐⭐⭐'],
    ['10', 'Xu và cs.\n(2021)', 'Trung Quốc\n(Hà Nam)', '373.216', 'GAD-7', 'Lo âu: 9,89%', 'CỠ MẪU LỚN NHẤT toàn cầu. Nam > nữ (10,1% vs 9,7%). Nông thôn cao nhất (12,8%). COVID worry/fear là yếu tố nguy cơ chính.', '⭐⭐⭐⭐⭐'],
    ['11', 'Bhardwaj và cs.\n(2020)', 'Ấn Độ\n(Chandigarh)', '288', 'DASS-21', 'Lo âu: 73,3%\nTrầm cảm: 64,9%\nStress: 74,7%', 'Tỷ lệ cao nhất 11 bài. Lo âu nặng+cực nặng 46,8%. Trường công lập, kinh tế thấp.', '⭐⭐'],
]

headers = ['#', 'Tác giả\n(Năm)', 'Quốc gia', 'n', 'Công cụ', 'Tỷ lệ', 'Phát hiện nổi bật', 'CL']
widths = [0.4, 1.6, 1.2, 0.9, 1.0, 1.6, 6.5, 0.5]

t = doc.add_table(rows=1+len(data), cols=len(headers))
t.style = 'Table Grid'
t.alignment = WD_TABLE_ALIGNMENT.CENTER

for row in t.rows:
    for ci in range(len(headers)):
        set_w(row.cells[ci], widths[ci])

# Header
for i, h in enumerate(headers):
    c = t.rows[0].cells[i]
    c.text = h
    for pp in c.paragraphs:
        pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in pp.runs:
            r.bold = True
            r.font.name = 'Times New Roman'
            r.font.size = Pt(7)
    shade(c, 'D9E2F3')

# Data
for ri, rd in enumerate(data):
    for ci, v in enumerate(rd):
        c = t.rows[ri+1].cells[ci]
        c.text = str(v)
        for pp in c.paragraphs:
            for r in pp.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(7)

doc.add_paragraph()

# TỔNG HỢP NHANH
p2 = doc.add_paragraph()
r2 = p2.add_run('TỔNG HỢP NHANH')
r2.bold = True
r2.font.name = 'Times New Roman'
r2.font.size = Pt(10)

summaries = [
    ('Tỷ lệ lo âu:', '9,89% (Xu, TQ n=373K, GAD-7) đến 73,3% (Bhardwaj, Ấn Độ, DASS-21). Dao động do công cụ, ngưỡng cắt, mẫu, bối cảnh. Chẩn đoán (NSCH Mỹ): 16,1%.'),
    ('Giới tính:', '8/11 bài: nữ > nam. Ngoại lệ: Saikia (nam>nữ ở Ấn Độ), Wen (nam>nữ nông thôn TQ), Xu (nam>nữ COVID TQ).'),
    ('Yếu tố nguy cơ:', 'Giấc ngủ kém (Chen, Mandaknalli), bắt nạt (Chen), sử dụng chất (Nakie AOR=5,6), nuôi dạy tiêu cực (Qiu OR=2,01), COVID worry (Xu), gia đình đơn thân (Saikia).'),
    ('Yếu tố bảo vệ:', 'Hỗ trợ tâm thần tại trường (Wen), nuôi dạy tích cực (Qiu OR=0,30), phục hồi cao (Qiu), hoạt động thể chất (Mandaknalli).'),
    ('Đáng chú ý:', 'Xu 2021 (n=373.216) lớn nhất toàn cầu. NSCH (n=55.162) lo âu tăng 61% trong 7 năm. Chen (n=63.205) tiên phong miền Tây TQ.'),
    ('Gap chung:', 'Thiếu RCT can thiệp trường học (đặc biệt châu Á). Thiếu NC dọc. Cần chuẩn hóa công cụ đo liên quốc gia.'),
]

for label, text in summaries:
    p = doc.add_paragraph()
    r = p.add_run(label + ' ')
    r.bold = True
    r.font.name = 'Times New Roman'
    r.font.size = Pt(8)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(8)

# CL legend
p3 = doc.add_paragraph()
r3 = p3.add_run('CL = Chất lượng: ⭐⭐ Thấp | ⭐⭐⭐ Trung bình | ⭐⭐⭐⭐ Tốt | ⭐⭐⭐⭐⭐ Xuất sắc')
r3.font.name = 'Times New Roman'
r3.font.size = Pt(7)
r3.italic = True

fname = 'Tóm tắt 2 trang - 11 bài nghiên cứu.docx'
doc.save(fname)
