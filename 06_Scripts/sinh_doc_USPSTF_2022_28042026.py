# -*- coding: utf-8 -*-
"""Doc trả lời thầy: USPSTF 2022 là gì?
Tiếng Việt thuần + chú thích Anh + reference DOI/PMID đầy đủ.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\USPSTF_2022_giai_thich_cho_thay_28042026.docx'
RED  = RGBColor(0xC0, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55); GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def warn(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run('⚠ '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def vn_apply(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.4)
    r=p.add_run('[Áp dụng cho Việt Nam] '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(12); r2.font.name='Times New Roman'

# =====================================================================
title("USPSTF 2022 LÀ GÌ?", 17)
subtitle("Trả lời thầy về thuật ngữ thường gặp trong tài liệu sàng lọc sức khỏe tâm thần trẻ em + vị thành niên")
nr("")
subtitle("Trợ lý nghiên cứu — 28/04/2026 — tiếng Việt thuần — chú thích thuật ngữ tiếng Anh trong ngoặc")

# =====================================================================
H1("1. CÂU TRẢ LỜI NHANH")
nr("USPSTF = U.S. Preventive Services Task Force = Lực lượng Đặc nhiệm Dịch vụ Phòng "
   "ngừa Hoa Kỳ. Đây là HỘI ĐỒNG CHUYÊN GIA ĐỘC LẬP do Bộ Y tế Hoa Kỳ thành lập từ "
   "năm 1984, có nhiệm vụ rà soát bằng chứng và đưa ra các khuyến cáo về SÀNG LỌC, "
   "ĐỀ PHÒNG, TƯ VẤN trong chăm sóc y tế ban đầu (primary care).", bold=True)
nr("")
nr("\"USPSTF 2022\" trong bối cảnh sức khỏe tâm thần trẻ em + vị thành niên thường chỉ "
   "BA KHUYẾN CÁO LỊCH SỬ được USPSTF công bố đồng thời tháng 10/2022 trên tạp chí "
   "JAMA:", bold=True)
nr("")
nr("(1) Sàng lọc rối loạn LO ÂU ở trẻ em + vị thành niên 8–18 tuổi — Mức B "
   "(KHUYẾN CÁO LÀM) — đây là LẦN ĐẦU TIÊN trong lịch sử USPSTF khuyến cáo sàng lọc "
   "lo âu cho nhóm tuổi này.", color=BLUE, bold=True)
nr("(2) Sàng lọc rối loạn TRẦM CẢM CHÍNH ở vị thành niên 12–18 tuổi — Mức B "
   "(KHUYẾN CÁO LÀM); ở trẻ em ≤11 tuổi — Mức I (BẰNG CHỨNG CHƯA ĐỦ).", color=BLUE, bold=True)
nr("(3) Sàng lọc NGUY CƠ TỰ TỬ ở trẻ em + vị thành niên — Mức I (BẰNG CHỨNG CHƯA ĐỦ).",
   color=BLUE, bold=True)

# =====================================================================
H1("2. USPSTF LÀ TỔ CHỨC GÌ?")

H2("2.1. Định nghĩa và lịch sử")
nr("USPSTF (United States Preventive Services Task Force — Lực lượng Đặc nhiệm Dịch vụ "
   "Phòng ngừa Hoa Kỳ) là một hội đồng gồm 16 chuyên gia y tế ĐỘC LẬP, làm việc tình "
   "nguyện, không nhận thù lao.")
nr("• Thành lập: năm 1984, do Cơ quan Nghiên cứu và Chất lượng Y tế Hoa Kỳ (Agency for "
   "Healthcare Research and Quality — AHRQ) tài trợ + hỗ trợ vận hành")
nr("• Cơ quan tài trợ: AHRQ, thuộc Bộ Y tế và Dịch vụ Nhân sinh Hoa Kỳ (U.S. Department "
   "of Health and Human Services — HHS)")
nr("• Thành viên: 16 chuyên gia trong các lĩnh vực: y học gia đình, nhi khoa, nội khoa, "
   "sản phụ khoa, y học dự phòng, dịch tễ học, chăm sóc dựa trên bằng chứng, y học hành "
   "vi, dược lý lâm sàng, kinh tế y tế")
nr("• Nhiệm kỳ thành viên: 4 năm; có thể tái nhiệm")
nr("• Trang chính thức: https://www.uspreventiveservicestaskforce.org/")

H2("2.2. Vai trò + tầm ảnh hưởng")
nr("USPSTF rà soát một cách HỆ THỐNG các bằng chứng khoa học có sẵn (tổng quan hệ "
   "thống, phân tích gộp, thử nghiệm có đối chứng) để đưa ra khuyến cáo về:")
nr("• Sàng lọc bệnh (screening tests)")
nr("• Tư vấn hành vi (behavioral counseling)")
nr("• Phòng ngừa bằng thuốc (preventive medications)")
nr("Khuyến cáo USPSTF có ẢNH HƯỞNG LỚN đến chính sách y tế Hoa Kỳ vì:")
nr("• Theo Đạo luật Chăm sóc Giá phải chăng (Affordable Care Act — ACA) năm 2010, các "
   "công ty bảo hiểm BẮT BUỘC chi trả 100% (không co-pay) cho các dịch vụ phòng ngừa "
   "được USPSTF xếp Mức A hoặc B.")
nr("• Khuyến cáo USPSTF được các tổ chức chuyên môn tham khảo + tích hợp vào hướng dẫn "
   "lâm sàng riêng (vd: Hiệp hội Nhi khoa Hoa Kỳ — American Academy of Pediatrics — AAP).")

H2("2.3. Hệ thống xếp hạng khuyến cáo (Grading System)")
nr("USPSTF xếp khuyến cáo theo 5 mức A / B / C / D / I:", bold=True)
tbl(['Mức', 'Ý nghĩa', 'Hành động khuyến nghị', 'Hệ quả bảo hiểm Hoa Kỳ'],
    [
        ['A', 'Bằng chứng MẠNH cho lợi ích RÕ',
         'BẮT BUỘC làm — cung cấp/giới thiệu dịch vụ',
         'Bảo hiểm BẮT BUỘC chi trả 100%'],
        ['B', 'Bằng chứng cho lợi ích VỪA → VỪA / hoặc lợi ích NHỎ → MẠNH',
         'NÊN làm — khuyến nghị cung cấp/giới thiệu',
         'Bảo hiểm BẮT BUỘC chi trả 100%'],
        ['C', 'Lợi ích NHỎ — chỉ làm khi có chỉ định cá nhân hoá',
         'CÂN NHẮC — quyết định cá nhân hoá theo bệnh nhân',
         'Bảo hiểm có thể chi trả nhưng không bắt buộc'],
        ['D', 'Bằng chứng KHÔNG có lợi ích / hoặc CÓ HẠI nhiều hơn lợi ích',
         'KHÔNG khuyến nghị — không cung cấp',
         'Không bắt buộc'],
        ['I', 'Bằng chứng CHƯA ĐỦ — không thể kết luận lợi/hại',
         'CÂN NHẮC bối cảnh; không khuyến nghị + cũng không phản đối',
         'Không bắt buộc'],
    ], [1.0, 5.0, 5.0, 5.0])

# =====================================================================
H1("3. KHUYẾN CÁO USPSTF 2022 — CHI TIẾT 3 BÀI")

H2("3.1. Sàng lọc Lo âu ở trẻ em + vị thành niên (Khuyến cáo Mức B — LỊCH SỬ)")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài', '\"Screening for Anxiety in Children and Adolescents: '
         'US Preventive Services Task Force Recommendation Statement\"'],
        ['Tạm dịch', 'Sàng lọc lo âu ở trẻ em và vị thành niên: Tuyên bố khuyến cáo '
         'của Lực lượng Đặc nhiệm Dịch vụ Phòng ngừa Hoa Kỳ'],
        ['Tác giả', 'US Preventive Services Task Force (Mangione CM, Barry MJ, '
         'Nicholson WK, Cabana M, Chelmow D, Coker TR và cộng sự — '
         'tổng cộng 16 thành viên)'],
        ['Tạp chí', 'JAMA — Journal of the American Medical Association'],
        ['Năm + ngày', '2022 — đăng ngày 11/10/2022'],
        ['Tập / Số / Trang', 'Tập 328(14): trang 1438–1444'],
        ['DOI', '10.1001/jama.2022.16936'],
        ['PMID', '36219403'],
        ['Đối tượng khuyến cáo', 'Trẻ em + vị thành niên 8-18 tuổi (KHÔNG có triệu chứng — '
         'sàng lọc đại chúng phòng ngừa)'],
        ['Khuyến cáo', 'MỨC B — \"USPSTF khuyến cáo sàng lọc rối loạn lo âu ở trẻ em + '
         'vị thành niên 8-18 tuổi\"'],
        ['Tính lịch sử', 'ĐÂY LÀ LẦN ĐẦU TIÊN trong lịch sử USPSTF khuyến cáo sàng lọc '
         'lo âu ở nhóm tuổi này — trước đó (2016) đã đưa ra khuyến cáo Mức I (chưa đủ '
         'bằng chứng)'],
    ], [4.0, 12.0])
nr("")
nr("Ý NGHĨA QUAN TRỌNG:", bold=True, color=BLUE)
nr("(1) Trước 2022, không có khuyến cáo chính thức nào về sàng lọc lo âu ở trẻ em — "
   "lo âu thường bị bỏ sót dù là rối loạn tâm thần phổ biến nhất ở nhóm này.")
nr("(2) Sau khuyến cáo, BẢO HIỂM HOA KỲ BẮT BUỘC CHI TRẢ 100% cho việc sàng lọc lo âu "
   "ở trẻ 8-18 tuổi — gỡ bỏ rào cản tài chính.")
nr("(3) Khuyến cáo dựa trên tổng quan hệ thống bằng chứng (evidence review) bởi Trung "
   "tâm Bằng chứng Khoa học EPC — cập nhật trên tổng quan trước đó năm 2016.")
nr("(4) Công cụ sàng lọc gợi ý: SCARED, GAD-7, Spence, PROMIS Anxiety Short Form — "
   "USPSTF không chỉ định công cụ cụ thể, để bác sĩ lâm sàng chọn theo bối cảnh.")

H2("3.2. Sàng lọc Trầm cảm chính ở trẻ em + vị thành niên (Mức B cho 12-18 / Mức I cho ≤11)")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài', '\"Screening for Depression and Suicide Risk in Children and '
         'Adolescents: US Preventive Services Task Force Recommendation Statement\"'],
        ['Tạm dịch', 'Sàng lọc trầm cảm và nguy cơ tự tử ở trẻ em + vị thành niên: '
         'Tuyên bố khuyến cáo USPSTF'],
        ['Tạp chí', 'JAMA'],
        ['Năm + ngày', '2022 — đăng ngày 11/10/2022 (cùng số với khuyến cáo lo âu)'],
        ['Tập / Số / Trang', 'Tập 328(15): trang 1534–1542'],
        ['DOI', '10.1001/jama.2022.16946'],
        ['PMID', '36219440'],
        ['Khuyến cáo về trầm cảm chính (MDD)',
         'MỨC B cho VỊ THÀNH NIÊN 12-18 tuổi (khuyến nghị sàng lọc); '
         'MỨC I cho TRẺ EM ≤11 tuổi (bằng chứng chưa đủ)'],
        ['Khuyến cáo về tự tử',
         'MỨC I cho TẤT CẢ trẻ em + vị thành niên (bằng chứng chưa đủ để khuyến cáo '
         'sàng lọc đại chúng nguy cơ tự tử ở nhóm không có triệu chứng)'],
    ], [4.0, 12.0])
nr("")
nr("LƯU Ý QUAN TRỌNG:", bold=True, color=BLUE)
nr("(1) Khuyến cáo MỨC B cho sàng lọc trầm cảm vị thành niên đã có từ 2016 — nhắc lại "
   "+ củng cố trong 2022.")
nr("(2) Mức I (chưa đủ bằng chứng) cho sàng lọc tự tử KHÔNG phải nói \"tự tử không quan "
   "trọng\" — mà là chưa đủ bằng chứng RẰNG SÀNG LỌC ĐẠI CHÚNG (không có triệu chứng) "
   "có lợi hơn hại. USPSTF vẫn khuyến cáo CAN THIỆP với người ĐÃ có dấu hiệu nguy cơ.")
nr("(3) Công cụ sàng lọc gợi ý cho trầm cảm: PHQ-9 (Patient Health Questionnaire-9), "
   "PHQ-A (phiên bản vị thành niên), MFQ, Beck Depression Inventory.")

H2("3.3. Tóm gọn 3 khuyến cáo USPSTF tháng 10/2022 (cho trẻ em + vị thành niên)")
tbl(['Vấn đề SKTT', 'Đối tượng', 'Mức USPSTF', 'Hệ quả thực tiễn'],
    [
        ['LO ÂU', '8-18 tuổi', 'B (KHUYẾN CÁO LÀM)',
         'Lần đầu — bảo hiểm bắt buộc chi trả'],
        ['TRẦM CẢM CHÍNH', '12-18 tuổi', 'B (KHUYẾN CÁO LÀM)',
         'Tiếp nối khuyến cáo 2016 — bảo hiểm bắt buộc chi trả'],
        ['TRẦM CẢM CHÍNH', '≤11 tuổi', 'I (CHƯA ĐỦ BẰNG CHỨNG)',
         'Không bắt buộc; bác sĩ tự cân nhắc'],
        ['NGUY CƠ TỰ TỬ', 'Tất cả tuổi <18', 'I (CHƯA ĐỦ BẰNG CHỨNG)',
         'Không khuyến cáo sàng lọc đại chúng; vẫn can thiệp khi có dấu hiệu'],
    ], [3.5, 3.0, 3.5, 6.0])

# =====================================================================
H1("4. TẠI SAO USPSTF 2022 LÀ DẤU MỐC LỊCH SỬ?")

H2("4.1. Bối cảnh trước 2022")
nr("Trước tháng 10/2022, lo âu là rối loạn tâm thần PHỔ BIẾN NHẤT ở trẻ em + vị thành "
   "niên Hoa Kỳ (theo CDC, ~9-10% trẻ em + vị thành niên có rối loạn lo âu được chẩn "
   "đoán) NHƯNG không có khuyến cáo sàng lọc chính thức nào ở cấp quốc gia. Hậu quả:")
nr("• Lo âu thường bị BỎ SÓT trong chăm sóc y tế ban đầu — chỉ ~1/3 trẻ có lo âu được "
   "chẩn đoán đúng + điều trị.")
nr("• Bảo hiểm có thể từ chối chi trả sàng lọc lo âu — gây rào cản tài chính.")
nr("• Bác sĩ không có hướng dẫn rõ ràng về CÔNG CỤ + TẦN SUẤT sàng lọc.")

H2("4.2. Tác động của khuyến cáo 2022")
nr("Sau khi USPSTF công bố Mức B tháng 10/2022:")
nr("(1) Bảo hiểm Hoa Kỳ BẮT BUỘC chi trả 100% sàng lọc lo âu ở trẻ 8-18 tuổi — gỡ "
   "rào cản tài chính.")
nr("(2) Hiệp hội Nhi khoa Hoa Kỳ (American Academy of Pediatrics — AAP) đã tích hợp "
   "vào hướng dẫn Bright Futures + khuyến cáo bác sĩ nhi khoa sàng lọc lo âu định kỳ.")
nr("(3) Nhiều hệ thống y tế đã triển khai sàng lọc lo âu trong khám sức khỏe định kỳ — "
   "kết quả: phát hiện sớm hơn, can thiệp kịp thời hơn.")
nr("(4) Khuyến cáo này TRỞ THÀNH chuẩn tham chiếu quốc tế — nhiều nước châu Âu, Úc, "
   "Canada cũng tham khảo khi xây dựng chính sách quốc gia.")

H2("4.3. Phản biện đối với khuyến cáo")
warn("Một số chuyên gia đã đưa ra các quan ngại:")
nr("(1) RỦI RO CHẨN ĐOÁN QUÁ MỨC (overdiagnosis): sàng lọc đại chúng có thể phát hiện "
   "nhiều \"trường hợp\" lo âu không thực sự cần điều trị — gây lo lắng cho gia đình + "
   "tốn nguồn lực điều trị.")
nr("(2) THIẾU NGUỒN LỰC ĐIỀU TRỊ: nếu sàng lọc nhưng không có dịch vụ điều trị đủ → "
   "tạo \"danh sách chờ\" + thất vọng. Hoa Kỳ thiếu trầm trọng nhà tâm lý lâm sàng "
   "chuyên về trẻ em — chỉ ~1/3 trẻ được chẩn đoán có thể tiếp cận điều trị.")
nr("(3) DỰ BÁO GIÁ TRỊ DƯƠNG TÍNH (positive predictive value) THẤP: ở quần thể tỉ lệ "
   "lo âu thấp, sàng lọc đại chúng cho nhiều DƯƠNG TÍNH GIẢ.")
nr("(4) THIẾU ĐA DẠNG VĂN HOÁ: các công cụ sàng lọc được phát triển chủ yếu ở bối cảnh "
   "Anglo-Saxon — có thể không phù hợp với nhóm thiểu số sắc tộc / di cư.")

# =====================================================================
H1("5. SO SÁNH USPSTF VỚI CÁC TỔ CHỨC KHÁC")
tbl(['Tổ chức', 'Quốc gia/Khu vực', 'Khuyến cáo về sàng lọc lo âu trẻ em + vị thành niên'],
    [
        ['USPSTF', 'Hoa Kỳ',
         'Mức B (8-18 tuổi) — khuyến cáo sàng lọc đại chúng (10/2022)'],
        ['American Academy of Pediatrics — AAP', 'Hoa Kỳ',
         'Tích hợp khuyến cáo USPSTF vào hướng dẫn Bright Futures cho khám định kỳ'],
        ['NICE (National Institute for Health and Care Excellence)', 'Anh',
         'Không khuyến cáo sàng lọc đại chúng cho nhóm KHÔNG triệu chứng; '
         'khuyến cáo đánh giá khi có dấu hiệu lâm sàng'],
        ['WHO (World Health Organization)', 'Toàn cầu',
         'Khuyến cáo tích hợp dịch vụ SKTT vào chăm sóc ban đầu (mhGAP); '
         'không có khuyến cáo cụ thể về sàng lọc đại chúng cho lo âu trẻ em'],
        ['Royal Australian and New Zealand College of Psychiatrists — RANZCP', 'Úc/NZ',
         'Khuyến cáo bác sĩ NHẬN DIỆN SỚM lo âu nhưng không khuyến cáo sàng lọc đại '
         'chúng có hệ thống'],
        ['Bộ Y tế Việt Nam', 'Việt Nam',
         'CHƯA có hướng dẫn quốc gia về sàng lọc lo âu trẻ em + vị thành niên trong '
         'chăm sóc ban đầu (đến 2026)'],
    ], [3.5, 2.5, 10.0])

# =====================================================================
H1("6. ÁP DỤNG CHO VIỆT NAM")
vn_apply("Khuyến cáo USPSTF 2022 có nhiều bài học cho Việt Nam:")
nr("(1) MÔ HÌNH CHÍNH SÁCH: Việt Nam có thể tham khảo khung USPSTF (rà soát hệ thống "
   "bằng chứng → khuyến cáo theo mức A/B/C/D/I) để xây dựng hướng dẫn quốc gia về "
   "sàng lọc SKTT trẻ em.")
nr("(2) ĐỐI TƯỢNG ƯU TIÊN: Mức B cho 8-18 tuổi phù hợp với học sinh THCS-THPT Việt Nam — "
   "giai đoạn áp lực học tập + thi đại học cao nhất.")
nr("(3) CÔNG CỤ SÀNG LỌC: GAD-7 đã có chuẩn Việt (Hoa 2024 với Cronbach α=0,916 ở Hà "
   "Nội) — có thể dùng làm công cụ chuẩn cho sàng lọc lo âu HS Việt Nam. SCARED chưa "
   "có bản Việt được kiểm định.")
nr("(4) HẠN CHẾ THỰC TIỄN: Trước khi triển khai sàng lọc đại chúng, cần đảm bảo có "
   "DỊCH VỤ ĐIỀU TRỊ tiếp nối — Việt Nam thiếu trầm trọng nhà tâm lý lâm sàng (chỉ ~"
   "0,2 chuyên gia SKTT/100.000 dân theo GBD 2021 ASEAN).")
nr("(5) MÔ HÌNH KHẢ THI: Tích hợp sàng lọc GAD-7 + PHQ-9 vào khám sức khỏe đầu năm "
   "học (đã có sẵn cơ sở pháp lý theo Thông tư 13/2016/TT-BYT về kiểm tra sức khỏe "
   "cho HS) — tận dụng hệ thống y tế trường học hiện có.")
nr("(6) BÀI HỌC PHẢN BIỆN: Cần TRÁNH chẩn đoán quá mức + có sẵn quy trình can thiệp "
   "cho trường hợp dương tính (chuyển tuyến đến tư vấn học đường + chuyên gia tâm lý "
   "liên trường).")

# =====================================================================
H1("7. ĐỐI CHIẾU VỚI KHO TÀI LIỆU DỰ ÁN")
nr("USPSTF 2022 có liên hệ với một số bài đã có trong dự án:")
nr("• Bài QT042_BESST (Brown 2024 Lancet Psychiatry — đã dịch + phản biện): nhấn mạnh "
   "tự giới thiệu (self-referral) như cơ chế thay thế cho sàng lọc đại chúng. Đây là "
   "GÓC NHÌN KHÁC với USPSTF — thay vì sàng lọc tất cả HS, để HS tự đăng ký.")
nr("• Bài QT042_PLACES (Brown 2022 IJERPH — đã dịch): mô hình PLACES với 6 yếu tố "
   "(Publicity, Lay, Acceptable, Convenient, Effective, Self-referral) là khung thiết "
   "kế dịch vụ giúp HS DỄ TIẾP CẬN sau khi được sàng lọc.")
nr("• Hoa 2024 (Hà Nội, n=3.910, GAD-7 với Cronbach α=0,916): xác nhận GAD-7 phù hợp "
   "Việt Nam — công cụ tiềm năng để áp dụng khuyến cáo USPSTF.")
nr("• V-NAMHS 2022 (Khảo sát quốc gia Việt Nam): cho thấy chỉ 8,4% VTN VN tiếp cận "
   "dịch vụ SKTT — Việt Nam cần tập trung vào TIẾP CẬN trước khi nói đến sàng lọc "
   "đại chúng.")

# =====================================================================
H1("8. THAM KHẢO ĐẦY ĐỦ")

H2("Bài chính USPSTF 2022")
nr("1. US Preventive Services Task Force, Mangione CM, Barry MJ, Nicholson WK và cộng "
   "sự (2022). Screening for Anxiety in Children and Adolescents: US Preventive "
   "Services Task Force Recommendation Statement. JAMA, 328(14):1438–1444. "
   "DOI: 10.1001/jama.2022.16936 — PMID 36219403", size=11)
nr("2. US Preventive Services Task Force (2022). Screening for Depression and Suicide "
   "Risk in Children and Adolescents: US Preventive Services Task Force Recommendation "
   "Statement. JAMA, 328(15):1534–1542. DOI: 10.1001/jama.2022.16946 — PMID 36219440",
   size=11)

H2("Bằng chứng nền (Evidence Reports được USPSTF dùng)")
nr("3. Viswanathan M, Wallace IF, Cook Middleton J và cộng sự (2022). Screening for "
   "Anxiety in Children and Adolescents: Evidence Report and Systematic Review for "
   "the US Preventive Services Task Force. JAMA, 328(14):1445–1455. "
   "DOI: 10.1001/jama.2022.16303", size=11)
nr("4. Viswanathan M, Middleton JC, Stuebe A và cộng sự (2022). Screening for "
   "Depression, Anxiety, and Suicide Risk in Children and Adolescents: Updated "
   "Evidence Report and Systematic Review for the US Preventive Services Task "
   "Force. JAMA, 328(15):1543–1556. DOI: 10.1001/jama.2022.16310", size=11)

H2("Bối cảnh + nguồn tham khảo về USPSTF")
nr("5. Trang chính thức USPSTF: https://www.uspreventiveservicestaskforce.org/", size=11)
nr("6. Bibbins-Domingo K, Grossman DC, Curry SJ và cộng sự (2016). Screening for "
   "Depression in Children and Adolescents: US Preventive Services Task Force "
   "Recommendation Statement. JAMA Pediatrics, 170(2):164–168. (Khuyến cáo cũ 2016 "
   "về trầm cảm — được nhắc lại + cập nhật năm 2022)", size=11)
nr("7. Hiệp hội Nhi khoa Hoa Kỳ (American Academy of Pediatrics — AAP) (2023). "
   "Bright Futures Guidelines for Health Supervision of Infants, Children, and "
   "Adolescents (phiên bản 4, cập nhật bao gồm khuyến cáo USPSTF 2022 về sàng lọc "
   "lo âu).", size=11)

H2("Bài liên quan trong kho dự án (đã dịch + phản biện)")
nr("• Brown JSL et al. (2024). BESST trial — Lancet Psychiatry 11(7):504-515. "
   "PMID 38759665 — corpus QT042_BESST", size=11)
nr("• Brown JSL et al. (2022). PLACES Model — IJERPH 19(5):2831. "
   "PMC8909998 — corpus QT042_PLACES", size=11)
nr("• Hoa P et al. (2024). GAD-7 chuẩn hoá ở HS THPT Hà Nội — corpus dự án", size=11)
nr("• V-NAMHS 2022 — Khảo sát quốc gia SKTT Việt Nam — corpus dự án", size=11)

d.save(OUT)
print('Wrote:', OUT)
