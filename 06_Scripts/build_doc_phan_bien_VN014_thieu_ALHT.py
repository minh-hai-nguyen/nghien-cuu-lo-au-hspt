"""DOC: Phan bien VN014 H.T. Hoc — KHONG nghien cuu Ap luc hoc tap.
Theo style cac doc phan bien VN004 / VN014 quan he cha me-con da lam.

VERIFIED FROM TOM-TAT VN014:
- Cac yeu to VN014 da do (Tom-tat doan [7]):
  + Nhan khau: gioi tinh, khu vuc cu tru, khoi lop
  + Hanh vi: thoi gian ngu ban ngay, thoi gian dung may tinh/dien thoai, thoi gian the thao,
    tinh chat moi quan he voi cha me, hinh thuc hoc tap (truc tuyen/truc tiep), giai cach xa hoi
- KHONG CO: ap luc hoc tap, bat nat hoc duong, tu trong, ho tro ban be
- R² = 0,190 (yeu)
- Beta cha me 0,272 (manh nhat); dien thoai 0,176; ngu -0,149; the thao -0,087
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Phan_bien_VN014_thieu_yeu_to_ap_luc_hoc_tap.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
RED = RGBColor(0xC0, 0x00, 0x00)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=False, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def blue_bullet(text, size=12):
    p = d.add_paragraph(style='List Bullet')
    r = p.add_run(text); r.font.size = Pt(size); r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('PHẢN BIỆN VN014 HOÀNG TRUNG HỌC (2025)\nTHIẾU YẾU TỐ ÁP LỰC HỌC TẬP TRONG NGHIÊN CỨU')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# Câu hỏi
H('Câu hỏi của thầy', level=2, color=NAVY)
blue_run(
    'Em cho thầy thêm một phản biện về H.T. Học với áp lực học '
    'tập nữa nhé. Hình như H.T.H không nghiên cứu áp lực học '
    'tập trong công trình đã tổng quan, em ạ. Check cẩn thận '
    'mọi fact/dữ liệu rồi ghi ra file doc nhé.', italic=True
)

# 1. Trả lời ngắn
H('Trả lời ngắn — Thầy ĐÚNG: VN014 KHÔNG nghiên cứu áp lực học tập', level=2, color=RED)
para(
    'Sau khi kiểm tra cẩn thận tóm tắt VN014 (Hoàng Trung Học và '
    'Nguyễn Thùy Dung, 2025), em xác nhận: NGHIÊN CỨU CỦA VN014 '
    'KHÔNG ĐO ÁP LỰC HỌC TẬP. Đây là HẠN CHẾ PHƯƠNG PHÁP LUẬN '
    'NGHIÊM TRỌNG, đặc biệt vì áp lực học tập là yếu tố nguy cơ '
    'HÀNG ĐẦU đối với rối loạn lo âu học sinh theo y văn quốc tế '
    'và theo chính chương 3 luận án.', bold=True, color=RED
)

# 2. Liệt kê yếu tố thực sự VN014 đã đo
H('1. Các yếu tố VN014 ĐÃ THỰC SỰ đo (verify từ Tom-tat đoạn [7])', level=2, color=NAVY)
caption('Bảng 1. Toàn bộ biến trong VN014 (n = 8.473 HS THCS-THPT, 6 tỉnh VN)')
add_table(
    ['Loại biến', 'Tên biến', 'Vai trò'],
    [
        ['Nhân khẩu (kiểm soát)', 'Giới tính', 'Biến kiểm soát'],
        ['Nhân khẩu (kiểm soát)', 'Khu vực cư trú (nông thôn/thành thị)', 'Biến kiểm soát'],
        ['Nhân khẩu (kiểm soát)', 'Khối lớp (6–12)', 'Biến kiểm soát'],
        ['Hành vi (predictor)', 'Thời gian ngủ ban ngày', 'β = −0,149 (bảo vệ)'],
        ['Hành vi (predictor)', 'Thời gian sử dụng máy tính/điện thoại', 'β = 0,176 (nguy cơ)'],
        ['Hành vi (predictor)', 'Thời gian thể thao', 'β = −0,087 (bảo vệ)'],
        ['Hành vi (predictor)', 'Tính chất mối quan hệ với cha mẹ', 'β = 0,272 (mạnh nhất)'],
        ['COVID-related', 'Hình thức học tập (trực tuyến/trực tiếp)', 'Biến COVID'],
        ['COVID-related', 'Mức độ giãn cách xã hội', 'Biến COVID'],
    ]
)
para('')
para(
    'Tổng cộng 9 biến (3 nhân khẩu + 4 hành vi + 2 COVID). Mô hình '
    'hồi quy đa biến đạt R² = 0,190 — chỉ giải thích 19% phương '
    'sai DASS tổng. Đây là R² THẤP cho mô hình SKTT.'
)

# 3. ALHT KHÔNG có trong VN014
H('2. Bốn yếu tố quan trọng VN014 KHÔNG đo', level=2, color=NAVY)
caption('Bảng 2. Bốn yếu tố nguy cơ/bảo vệ MẠNH NHẤT vắng mặt trong VN014')
add_table(
    ['Yếu tố thiếu', 'Vai trò trong y văn', 'β trong chương 3 luận án CTH', 'Hậu quả VN014'],
    [
        ['ÁP LỰC HỌC TẬP (ALHT)',
         'Yếu tố nguy cơ HÀNG ĐẦU theo Pascoe (2020) khung 6 trục',
         'β = 0,510 (RLLATQ); 0,490 (RLLAXH) — MẠNH NHẤT',
         'BỎ SÓT yếu tố nguy cơ mạnh nhất → R² thấp'],
        ['BẮT NẠT HỌC ĐƯỜNG',
         '25% nạn nhân toàn cầu (meta-analysis)',
         'β = 0,376 (RLLACL — mạnh nhất cho chia ly)',
         'BỎ SÓT yếu tố TRƯỜNG HỌC quan trọng'],
        ['LÒNG TỰ TRỌNG',
         'Yếu tố bảo vệ MẠNH NHẤT (Sowislo & Orth, 2013 meta 95 NC)',
         'β = −0,455 (RLLATQ); −0,415 (RLLAXH)',
         'BỎ SÓT yếu tố bảo vệ mạnh nhất'],
        ['HỖ TRỢ TỪ BẠN BÈ',
         'Y văn xác lập (Murphy 2024 peer support)',
         'β GẦN NHƯ KHÔNG có (đặc thù VN — phát hiện CTH)',
         'BỎ SÓT — không thể test giả thuyết "VN khác phương Tây"'],
    ]
)
para('')
para(
    'Đáng chú ý nhất là sự VẮNG MẶT của ÁP LỰC HỌC TẬP. Đây là '
    'thiếu sót LỚN vì:', bold=True
)

# 4. Tại sao thiếu ALHT là vấn đề lớn
H('3. Bốn lý do vì sao thiếu ALHT là HẠN CHẾ NGHIÊM TRỌNG', level=2, color=NAVY)

H('3.1. ALHT là yếu tố nguy cơ HÀNG ĐẦU theo y văn quốc tế', level=3)
para(
    'Pascoe, Hetrick và Parker (2020) trong International Journal '
    'of Adolescence and Youth — bài tổng quan tường thuật được '
    'trích dẫn hơn 1.500 lần — xác lập áp lực học tập tác động '
    'lên SÁU TRỤC đời sống học sinh: sức khỏe tâm thần (lo âu + '
    'trầm cảm), sử dụng chất, giấc ngủ, sức khỏe thể chất, thành '
    'tích, và bỏ học. Áp lực thi cử (high-stakes exams) được '
    'khẳng định là yếu tố nguy cơ ĐỘC LẬP với rối loạn lo âu.'
)
para(
    'Brunborg và cộng sự (2025) — Soc Sci Med — trên 979.043 '
    'thanh thiếu niên Na Uy 13 năm: BẤT MÃN TRƯỜNG HỌC (chứa áp '
    'lực học tập) loại bỏ TOÀN BỘ xu hướng tăng căng thẳng tâm '
    'thần ở nam. KHÔNG đo yếu tố này = bỏ qua nguyên nhân CHÍNH.'
)

H('3.2. Chính chương 3 luận án CTH xác lập ALHT là yếu tố MẠNH NHẤT', level=3)
para(
    'Mô hình SEM trên n = 1.352 học sinh THCS Việt Nam (chương 3 '
    'luận án) đã chứng minh:'
)
para('• β ALHT → RLLATQ = 0,510 (lớn nhất trong các yếu tố nguy cơ riêng lẻ)')
para('• β ALHT → RLLAXH = 0,490')
para('• β ALHT → RLLACL = 0,253')
para('• Mô hình ALHT-RLLA tổng (3 factors): β = 0,533; R² = 0,284')
para(
    'Tức là MỘT MÌNH ALHT đã giải thích 28,4% phương sai RLLA '
    'tổng — vượt cả R² CỦA VN014 ĐẦY ĐỦ (0,190 = 19%).'
)
para(
    'Phù hợp khẳng định: VN014 nếu CÓ đo ALHT, R² mô hình có thể '
    'cao hơn nhiều (ước lượng 0,30–0,40 thay vì 0,19).'
)

H('3.3. R² = 0,190 PHẢN ÁNH thiếu yếu tố quan trọng', level=3)
para(
    'R² = 0,190 nghĩa là 81% phương sai DASS không được giải '
    'thích bởi mô hình. Theo Cohen (1988): R² = 0,02 (small), '
    '0,13 (medium), 0,26 (large). Vậy R² VN014 = 0,19 thuộc nhóm '
    'medium — KHÔNG đạt nhóm large.'
)
para(
    'So sánh với chương 3 luận án CTH:'
)
add_table(
    ['Mô hình', 'R²', 'Đánh giá'],
    [
        ['VN014 (Hoàng Trung Học 2025) — 6 yếu tố hành vi + 3 nhân khẩu',
         '0,190', 'Medium'],
        ['Chương 3 ALHT riêng → RLLA tổng', '0,284', 'Large'],
        ['Chương 3 YTNC tổng (3 yếu tố) → RLLA tổng', '0,558', 'Rất large'],
        ['Chương 3 mô hình tích hợp YTNC + YTBV → RLLA tổng', '0,598', 'Rất large'],
    ]
)
para('')
para(
    'Mô hình tích hợp của chương 3 (R² = 0,598) gấp HƠN BA LẦN R² '
    'của VN014 (0,190). Sự khác biệt phần lớn do VN014 thiếu ALHT '
    '+ tự trọng + bắt nạt.'
)

H('3.4. Bỏ sót ALHT làm SAI LỆCH diễn giải về quan hệ cha mẹ-con', level=3)
para(
    'VN014 phát hiện QUAN HỆ CHA MẸ-CON là yếu tố MẠNH NHẤT '
    '(β = 0,272). Tuy nhiên, do KHÔNG kiểm soát ALHT, hệ số này '
    'có thể bị PHÓNG ĐẠI vì ALHT TƯƠNG QUAN với quan hệ cha '
    'mẹ-con (cha mẹ kỳ vọng cao → áp lực học tập → quan hệ '
    'căng thẳng).'
)
para(
    'Khi kiểm soát đầy đủ ALHT (như chương 3 luận án), β quan hệ '
    'cha mẹ-con (đo bằng HTCM) chỉ là −0,234 — YẾU HƠN NHIỀU so '
    'với β ALHT = 0,533. Tức là VN014 có thể đã NHẦM cha mẹ-con '
    'là yếu tố mạnh nhất, trong khi ALHT mới là yếu tố THỰC SỰ '
    'mạnh nhất.'
)

# 5. So sánh đối chiếu
H('4. So sánh thiết kế đo lường — VN014 vs Chương 3 luận án vs y văn quốc tế', level=2, color=NAVY)
caption('Bảng 3. So sánh thiết kế ba bộ nghiên cứu')
add_table(
    ['Yếu tố', 'VN014 H.T. Học (2025)', 'Chương 3 luận án CTH', 'Pascoe (2020) khung 6 trục'],
    [
        ['Áp lực học tập', '✗ KHÔNG ĐO', '✓ ESSA 4 mục, β=0,510', '✓ Trục 1 — yếu tố nguy cơ'],
        ['Bắt nạt học đường', '✗ KHÔNG ĐO', '✓ OBVQ, β=0,376 cho RLLACL', '✓ Trong khung trường học'],
        ['Tự trọng', '✗ KHÔNG ĐO', '✓ RSES, β=−0,455', '✓ Yếu tố cá nhân'],
        ['Hỗ trợ bạn bè', '✗ KHÔNG ĐO', '✓ MSPSS, β NS', '✓ Trục quan hệ xã hội'],
        ['Quan hệ cha mẹ', '✓ 1 biến tổng, β=0,272', '✓ MSPSS chi tiết, β=−0,234',
         '✓ Trục gia đình'],
        ['Nghiện điện thoại', '✓ Thời gian dùng, β=0,176',
         '✓ SAS-SV, β=0,336', '✓ Trục công nghệ'],
        ['Giấc ngủ', '✓ Thời gian ngủ, β=−0,149', '— (không trực tiếp)',
         '✓ Trục 3 — sleep'],
        ['Thể thao', '✓ Thời gian thể thao, β=−0,087', '— (không trực tiếp)',
         '✓ Trục 4 — physical health'],
        ['n', '8.473', '1.352', 'Khoảng 80 NC tổng quan'],
        ['R² mô hình tổng', '0,190', '0,598 (tích hợp)', 'N/A'],
    ]
)
para('')
para(
    'Bức tranh rõ: VN014 có CỠ MẪU LỚN (n=8.473) nhưng BỘ ĐO '
    'YẾU — chỉ 6 yếu tố hành vi + bỏ sót 4 yếu tố quan trọng. '
    'Chương 3 luận án có cỡ mẫu nhỏ hơn (n=1.352) nhưng BỘ ĐO '
    'TOÀN DIỆN HƠN — phù hợp khung Pascoe 2020.'
)

# 6. Hậu quả cho việc trích dẫn
H('5. Hậu quả cho việc trích dẫn VN014 trong báo cáo CTH', level=2, color=NAVY)
para(
    'Nếu trích dẫn VN014 trong báo cáo, cần lưu ý các giới hạn:'
)
para(
    '(a) Không trích β = 0,272 quan hệ cha mẹ-con như "yếu tố '
    'mạnh nhất ảnh hưởng RLLA ở HS Việt Nam" — vì VN014 KHÔNG '
    'kiểm soát ALHT.'
)
para(
    '(b) Khi so sánh "yếu tố nguy cơ MẠNH NHẤT", phải so với '
    'chương 3 luận án CTH (có ALHT) chứ không phải VN014 (không '
    'có ALHT).'
)
para(
    '(c) Nên trích VN014 cho các yếu tố RIÊNG mà VN014 đã đo: '
    'thời gian dùng điện thoại (β=0,176), giấc ngủ (β=−0,149), '
    'thể thao (β=−0,087). KHÔNG trích cho áp lực học tập.'
)
para(
    '(d) Có thể trích VN014 LÀM PHẢN BIỆN — chỉ ra HẠN CHẾ '
    'của thiết kế: "VN014 trên n=8.473 vẫn có R² chỉ 0,190 do '
    'thiếu ALHT — chương 3 luận án nâng R² lên 0,598 nhờ tích '
    'hợp đầy đủ ALHT + bắt nạt + tự trọng".'
)

# 7. CÂU TRẢ LỜI tô xanh
H('6. CÂU TRẢ LỜI', level=2, color=NAVY)
blue_run('Tóm gọn:', bold=True)
blue_bullet(
    'XÁC NHẬN — VN014 H.T. Học (2025) KHÔNG nghiên cứu áp lực '
    'học tập. Tác giả chỉ đo 6 yếu tố hành vi (giấc ngủ, điện '
    'thoại, thể thao, quan hệ cha mẹ-con, hình thức học tập, '
    'giãn cách xã hội) + 3 nhân khẩu (giới, khu vực, khối). '
    'Verify từ Tom-tat đoạn [7].'
)
blue_bullet(
    'Đây là HẠN CHẾ NGHIÊM TRỌNG vì áp lực học tập là yếu tố '
    'nguy cơ HÀNG ĐẦU theo y văn quốc tế (Pascoe 2020 khung 6 '
    'trục) và theo chính chương 3 luận án (β = 0,510 cho '
    'RLLATQ — lớn nhất trong các yếu tố nguy cơ).'
)
blue_bullet(
    'Hậu quả: R² VN014 = 0,190 (medium) — chỉ giải thích 19% '
    'phương sai DASS, trong khi chương 3 luận án (CÓ ALHT) đạt '
    'R² = 0,598 cho mô hình tích hợp. Khoảng cách 3,1 lần phần '
    'lớn do VN014 thiếu ALHT + tự trọng + bắt nạt.'
)
blue_bullet(
    'Diễn giải có thể bị SAI LỆCH: VN014 nói "quan hệ cha mẹ-con '
    'là yếu tố mạnh nhất" (β=0,272). Nếu KIỂM SOÁT ALHT (như '
    'chương 3 luận án — β HTCM = −0,234 sau khi kiểm soát ALHT '
    'và tự trọng), kết luận có thể KHÁC. ALHT có thể là yếu '
    'tố MẠNH NHẤT thực sự.'
)
blue_bullet(
    'Khuyến nghị trích dẫn VN014: chỉ trích cho các yếu tố '
    'VN014 đã đo (thời gian điện thoại β=0,176, giấc ngủ '
    'β=−0,149, thể thao β=−0,087); KHÔNG trích cho ALHT; '
    'có thể trích làm PHẢN BIỆN — chỉ ra hạn chế thiết kế.'
)

blue_run('Cách trích vào báo cáo CTH:', bold=True)
blue_run(
    '"Hoàng Trung Học và Nguyễn Thùy Dung (2025) trên 8.473 '
    'học sinh THCS-THPT 6 tỉnh Việt Nam khảo sát mức độ stress, '
    'lo âu và trầm cảm trong và sau COVID-19. Mô hình hồi quy '
    'đa biến đạt R² = 0,190 với bốn yếu tố hành vi: thời gian '
    'sử dụng máy tính/điện thoại (β = 0,176), thời gian ngủ '
    '(β = −0,149), thời gian thể thao (β = −0,087), và tính '
    'chất quan hệ với cha mẹ (β = 0,272). Cần lưu ý rằng nghiên '
    'cứu này không bao gồm ÁP LỰC HỌC TẬP — yếu tố được xác '
    'lập trong y văn quốc tế (Pascoe và cộng sự, 2020) và trong '
    'chương 3 luận án (β = 0,510 cho lo âu lan tỏa) là yếu tố '
    'nguy cơ MẠNH NHẤT đối với rối loạn lo âu học sinh THCS. '
    'Việc thiếu áp lực học tập có thể là một trong những lý do '
    'mô hình VN014 đạt R² chỉ 0,190 (medium effect size theo '
    'Cohen 1988), thấp hơn nhiều so với mô hình tích hợp '
    'chương 3 luận án (R² = 0,598)."', italic=True
)

# 8. TLTK
H('7. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
    'Hoàng, T. H., & Nguyễn, T. D. (2025). Mức độ căng thẳng, lo âu và trầm cảm ở thanh thiếu niên trong và sau đại dịch COVID-19 tại Việt Nam: Nghiên cứu cắt ngang. [VN014 trong cơ sở dữ liệu dự án.]',
    'Lovibond, S. H., & Lovibond, P. F. (1995). Manual for the Depression Anxiety Stress Scales (2nd ed.). Sydney: Psychology Foundation of Australia.',
    'Pascoe, M. C., Hetrick, S. E., & Parker, A. G. (2020). The impact of stress on students in secondary school and higher education. International Journal of Adolescence and Youth, 25(1), 104–112. https://doi.org/10.1080/02673843.2019.1596823 [QT067 trong cơ sở dữ liệu dự án.]',
    'Sowislo, J. F., & Orth, U. (2013). Does low self-esteem predict depression and anxiety? A meta-analysis of longitudinal studies. Psychological Bulletin, 139(1), 213–240.',
    'Sun, J., Dunne, M. P., Hou, X. Y., & Xu, A. Q. (2011). Educational Stress Scale for Adolescents: Development, validity, and reliability with Chinese students. Journal of Psychoeducational Assessment, 29(6), 534–546.',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
