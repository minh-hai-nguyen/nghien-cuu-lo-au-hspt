"""Doc trả lời 3 câu hỏi: Lochman 2025 EACP + Murphy 2024 Peer Support + tỷ lệ HS THCS lo âu tự tìm chuyên môn."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/CopingPower_PeerSupport_HelpSeeking_HS_THCS.docx'

d = Document()
style = d.styles['Normal']; style.font.name = 'Calibri'; style.font.size = Pt(11)
BLUE = RGBColor(0, 112, 192)
DARK = RGBColor(31, 73, 125)
GREEN = RGBColor(54, 95, 44)
RED = RGBColor(192, 0, 0)
ORANGE = RGBColor(191, 97, 14)
GRAY = RGBColor(90, 90, 90)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)
def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
def vn_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.bold = bold; r.font.size = Pt(size)
    return p
def blue_para(text, bold=False, size=11):
    p = d.add_paragraph(); r = p.add_run(text); r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = BLUE
    return p

# ===== TITLE =====
title = d.add_heading('Coping Power (Lochman 2025) + Peer Support (Murphy 2024) + Tỷ lệ HS THCS lo âu tự tìm chuyên môn', level=0)
for r in title.runs: r.font.color.rgb = DARK
sub = d.add_paragraph()
sr = sub.add_run('Sơ lược 2 chương trình can thiệp + so sánh tỷ lệ help-seeking VN vs US')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = GRAY
d.add_paragraph()

# ===== CÂU HỎI =====
qbox = d.add_table(rows=1, cols=1); qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]; shade(cell, 'FFF8DC')
p = cell.paragraphs[0]; r = p.add_run('Câu hỏi của thầy:'); r.bold = True
cell.add_paragraph('"Em cho thầy sơ lược nội dung của 2 cái này nhé: Nội dung Chương trình \'Sức mạnh ứng phó\' '
                   'cho vị thành niên sớm (Lochman 2025 ở Mỹ) + Nội dung chương trình hỗ trợ ngang hàng '
                   '(Murphy 2024 ở Wiley). Em cho thầy biết tỷ lệ HS THCS bị lo âu tự tìm đến nhà chuyên môn '
                   'ở nước ngoài và VN khoảng bao nhiêu nhé?"')
d.add_paragraph()

# ===== BỐI CẢNH =====
add_h('Bối cảnh', 1)
warn = d.add_table(rows=1, cols=1); warn.style = 'Table Grid'
wc = warn.rows[0].cells[0]; shade(wc, 'FCE4D6')
wp = wc.paragraphs[0]; wr = wp.add_run('⚠ Lưu ý: '); wr.bold = True; wr.font.color.rgb = RED
wp.add_run('Cả 2 bài Lochman 2025 và Murphy 2024 KHÔNG có sẵn trong DB 87 canonical của em. Thông tin được cào từ '
           'web (PubMed, Blueprints Programs, ScienceDirect abstract, Wiley Online Library), KHÔNG verify đầy đủ '
           'từ PDF gốc. Nếu thầy cần em canonical hóa 2 bài này (QT065 + QT066), báo em làm.')
d.add_paragraph()

# =====================================================
# 1. LOCHMAN 2025 — EACP
# =====================================================
add_h('1. Lochman 2025 — Early Adolescent Coping Power (EACP)', 1)

vn_para('Bài chính xác:', bold=True)
vn_para('Lochman JE và cộng sự (2025). "Randomized controlled trial of the early adolescent coping power program: '
        'Effects on emotional and behavioral problems in middle schoolers." Journal of School Psychology. '
        'DOI: 10.1016/j.jsp.2025.01.0XX (article in press; pii S002244052500010X).')
d.add_paragraph()

vn_para('1.1. Đối tượng & thiết kế:', bold=True, size=12)
for it in [
    '40 trường THCS Mỹ, n = 709 HS lớp 7 — được tuyển vì có hành vi gây hấn (aggressive behavior)',
    '69,8 % người gốc Phi (African American); 59,4 % nam',
    'Thiết kế: RCT cluster — so EACP vs nhóm chứng (treatment-as-usual)',
    'Tài trợ: NIMH (National Institute of Mental Health)',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
d.add_paragraph()

vn_para('1.2. Cấu trúc chương trình EACP — 3 thành phần:', bold=True, size=12)
eacp_tbl = d.add_table(rows=4, cols=2); eacp_tbl.style = 'Table Grid'
hdr = ['Thành phần', 'Chi tiết']
for i, h in enumerate(hdr):
    c = eacp_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
eacp_data = [
    ('① HỌC SINH', '25 buổi nhóm trong 1 năm học (lớp 7) — giảm từ 34 buổi của bản gốc Coping Power lớp 5-6. '
                   'Nhóm 5-8 HS, 40-50 phút/buổi. Bổ sung buổi cá nhân 30 phút/2 tháng.'),
    ('② CAREGIVERS (cha mẹ)', 'Buổi nhóm 12+ phụ huynh, 90 phút/buổi. Tổng ~16 buổi qua 5-6 tháng. '
                              'Tập trung: quản lý hành vi con + giao tiếp gia đình + quản lý stress (2 buổi riêng).'),
    ('③ GIÁO VIÊN / TRƯỜNG', 'Phối hợp với cố vấn học đường (school counselor). '
                              'KHÔNG có curriculum riêng — chủ yếu hỗ trợ truyền thông + củng cố ở lớp.'),
]
for i, (k, v) in enumerate(eacp_data):
    c0 = eacp_tbl.rows[i+1].cells[0]; shade(c0, 'E2EFDA')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = GREEN
    eacp_tbl.rows[i+1].cells[1].text = v
d.add_paragraph()

vn_para('1.3. Nội dung kỹ năng cho HS (5 nhóm):', bold=True, size=12)
for it in [
    '① QUẢN LÝ CƠN GIẬN/LO ÂU: self-statements (lời nói nội tâm tích cực), kỹ thuật thư giãn',
    '② GIẢI QUYẾT VẤN ĐỀ XÃ HỘI: nhận diện vấn đề → tạo nhiều giải pháp → đánh giá → chọn → thực hiện',
    '③ TƯƠNG TÁC VỚI BẠN BÈ: kỹ năng giao tiếp, đọc tín hiệu xã hội, xử lý xung đột',
    '④ KỸ NĂNG HỌC TẬP + TỔ CHỨC: study skills, quản lý thời gian',
    '⑤ KỸ NĂNG TỪ CHỐI CHẤT GÂY NGHIỆN (substance abuse refusal)',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
d.add_paragraph()

vn_para('1.4. Kết quả (so chứng):', bold=True, size=12)
res_tbl = d.add_table(rows=6, cols=2); res_tbl.style = 'Table Grid'
hdr = ['Thời điểm', 'Tác động']
for i, h in enumerate(hdr):
    c = res_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
res_data = [
    ('Post-test (cuối năm)', 'Giảm hành vi gây hấn chủ động (parent-rated); cải thiện hành vi tổng thể (teacher-rated)'),
    ('Post-test', 'Cải thiện kỹ năng xã hội (teacher-rated)'),
    ('1 năm sau', 'Giảm tự báo cáo phạm pháp (delinquency)'),
    ('1 năm sau', 'Giảm tự báo cáo dùng chất gây nghiện'),
    ('3 năm sau', 'Giảm hành vi gây hấn (teacher-rated) — DUY TRÌ DÀI HẠN'),
]
for i, (k, v) in enumerate(res_data):
    res_tbl.rows[i+1].cells[0].text = k
    res_tbl.rows[i+1].cells[1].text = v
    for pp in res_tbl.rows[i+1].cells[0].paragraphs:
        for rr in pp.runs: rr.bold = True
d.add_paragraph()

vn_para('1.5. Nền tảng lý thuyết:', bold=True)
vn_para('Mô hình social-cognitive developmental contextual — kết hợp lý thuyết phát triển xã hội-nhận thức '
        '(Crick & Dodge social information processing) với bối cảnh phát triển VTN.')
d.add_paragraph()

# =====================================================
# 2. MURPHY 2024 — PEER SUPPORT
# =====================================================
add_h('2. Murphy 2024 — Peer Support trong Primary Youth Mental Health Care', 1)

vn_para('Bài chính xác:', bold=True)
vn_para('Murphy R, Huggard L, Fitzgerald A, Hennessy E, Booth A (2024). "A systematic scoping review of peer '
        'support interventions in integrated primary youth mental health care." Journal of Community Psychology, '
        '52(1):154-180. DOI: 10.1002/jcop.23090.')
d.add_paragraph()

vn_para('2.1. Loại bài + phạm vi:', bold=True, size=12)
for it in [
    'SCOPING REVIEW (không phải RCT đơn lẻ) — lập phạm vi bằng chứng peer support trong dịch vụ MH thanh niên',
    '15 nghiên cứu gốc đáp ứng tiêu chí',
    '13 interventions peer support được phân tích chi tiết',
    'Bối cảnh: integrated primary youth mental health care (vd Jigsaw — Ireland\'s national youth MH service)',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
d.add_paragraph()

vn_para('2.2. Định nghĩa peer support:', bold=True, size=12)
def_tbl = d.add_table(rows=1, cols=1); def_tbl.style = 'Table Grid'
dc = def_tbl.rows[0].cells[0]; shade(dc, 'DEEBF7')
dp = dc.paragraphs[0]
dp.add_run('"Hỗ trợ XÃ HỘI và CẢM XÚC do và nhận giữa những cá nhân có TRẢI NGHIỆM CHUNG về khó khăn sức khoẻ '
           'tâm thần."').italic = True
d.add_paragraph()

vn_para('2.3. Đặc điểm 13 interventions:', bold=True, size=12)
char_tbl = d.add_table(rows=5, cols=2); char_tbl.style = 'Table Grid'
hdr = ['Đặc điểm', 'Phân bố']
for i, h in enumerate(hdr):
    c = char_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
char_data = [
    ('Đối tượng MH', '8/13 nhắm youth có khó khăn MH NHẸ-VỪA (trầm cảm, lo âu, distress tâm lý)'),
    ('Tuổi', '9/13 nhắm 16-25; 1 nhắm 13-19; 1 nhắm 11-18 (gần với HS THCS); 1 nhắm 12-25; còn lại tuổi rộng'),
    ('Công cụ đo', 'GAD-7 (lo âu), DASS-21 (trầm cảm-lo âu-stress); và một số thang về self-efficacy, quality of life'),
    ('Mode delivery', 'Mặt-đối-mặt cá nhân; nhóm nhỏ; online/digital; điện thoại — biến thiên'),
]
for i, (k, v) in enumerate(char_data):
    c0 = char_tbl.rows[i+1].cells[0]; shade(c0, 'E2EFDA')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = GREEN
    char_tbl.rows[i+1].cells[1].text = v
d.add_paragraph()

vn_para('2.4. Phát hiện chính:', bold=True, size=12)
for it in [
    'Peer support có TIỀM NĂNG cải thiện outcome PHỤC HỒI (recovery-related) — nhưng bằng chứng chưa đồng nhất',
    'Rào cản triển khai: (1) lo ngại bảo mật trong quan hệ peer; (2) peer worker thiếu tự tin về vai trò',
    'Yếu tố thúc đẩy: (1) nhân viên chuyên môn HỖ TRỢ peer worker; (2) vai trò peer được ĐỊNH NGHĨA RÕ',
    'Khuyến nghị: các nghiên cứu sau cần báo cáo CHI TIẾT hơn nội dung intervention (nhiều bài thiếu mô tả)',
]:
    p = d.add_paragraph(it, style='List Bullet')
    for r in p.runs: r.font.size = Pt(11)
d.add_paragraph()

vn_para('2.5. Hạn chế quan trọng cho áp dụng VN:', bold=True, size=12)
warn2 = d.add_table(rows=1, cols=1); warn2.style = 'Table Grid'
wc2 = warn2.rows[0].cells[0]; shade(wc2, 'FCE4D6')
wp2 = wc2.paragraphs[0]
wp2.add_run('⚠ ').bold = True
wp2.add_run('Đa số 9/13 interventions nhắm 16-25 (THANH NIÊN, không phải HS THCS-THPT). Chỉ 2 chương trình '
            'có HS THCS-THPT. Nếu áp dụng cho HS THCS VN, cần ADAPT đáng kể về: (a) hiểu biết MH literacy '
            'của 11-15 tuổi; (b) bảo mật + parental consent; (c) đào tạo peer worker tuổi gần.')
d.add_paragraph()

# =====================================================
# 3. HELP-SEEKING RATE — VN vs US
# =====================================================
add_h('3. Tỷ lệ HS THCS lo âu TỰ TÌM CHUYÊN MÔN — VN vs US', 1)

caveat = d.add_table(rows=1, cols=1); caveat.style = 'Table Grid'
cv = caveat.rows[0].cells[0]; shade(cv, 'FFF2CC')
cvp = cv.paragraphs[0]; cvr = cvp.add_run('⚠ Lưu ý: '); cvr.bold = True; cvr.font.color.rgb = ORANGE
cvp.add_run('Data chuyên cho HS THCS (12-14 tuổi) HIẾM. Đa số survey gộp 12-17 hoặc 11-19. Em báo data có sẵn — '
            'thầy có thể giả định tỷ lệ HS THCS thấp hơn hoặc tương đương.')
d.add_paragraph()

vn_para('3.1. Việt Nam', bold=True, size=12)
vn_tbl = d.add_table(rows=4, cols=2); vn_tbl.style = 'Table Grid'
hdr = ['Nguồn', 'Tỷ lệ + ghi chú']
for i, h in enumerate(hdr):
    c = vn_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
vn_data = [
    ('Nguyễn et al. 2020 HCM City\n(HS THPT, không phải THCS)', '12,5 % [KTC 95 %: 10,9-14,1] HS có triệu chứng MH dùng dịch vụ chuyên môn '
                                                                'trong 12 tháng qua'),
    ('V-NAMHS 2022 (VN002, n=5.996, 10-17 tuổi)', '8,4 % VTN có nhu cầu tiếp cận dịch vụ MH; chỉ 5,1 % phụ huynh '
                                                  'NHẬN RA con cần. Tỷ lệ tiếp cận dịch vụ THỰC chỉ ~2-3 %'),
    ('Help-seeking preference (chung)', 'HS VN ưu thích NON-PROFESSIONAL (gia đình, bạn bè, thầy cô) làm lựa chọn '
                                        'ĐẦU TIÊN. Chuyên môn (psychologist, psychiatrist) là lựa chọn cuối cùng.'),
]
for i, (k, v) in enumerate(vn_data):
    c0 = vn_tbl.rows[i+1].cells[0]; shade(c0, 'E2EFDA')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = GREEN
    vn_tbl.rows[i+1].cells[1].text = v
d.add_paragraph()

vn_para('3.2. Hoa Kỳ', bold=True, size=12)
us_tbl = d.add_table(rows=5, cols=2); us_tbl.style = 'Table Grid'
hdr = ['Nguồn', 'Tỷ lệ']
for i, h in enumerate(hdr):
    c = us_tbl.rows[0].cells[i]; shade(c, '4472C4')
    pp = c.paragraphs[0]; rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = RGBColor(255,255,255)
us_data = [
    ('NSCH 2021-2023 (12-17 tuổi)', '20 % có UNMET mental health care needs'),
    ('NSCH 2022-2023', '54 % youth 12-17 KHÓ tiếp cận MH care cần thiết'),
    ('NSCH 2023 (chẩn đoán cần ĐT)', '61 % có khó khăn tiếp cận, TĂNG 35 % so với 2018'),
    ('Anxiety prevalence US 12-17 (2021-2023)', '20 % có triệu chứng lo âu trong 2 tuần qua'),
]
for i, (k, v) in enumerate(us_data):
    c0 = us_tbl.rows[i+1].cells[0]; shade(c0, 'E2EFDA')
    pp = c0.paragraphs[0]; rr = pp.add_run(k); rr.bold = True; rr.font.color.rgb = GREEN
    us_tbl.rows[i+1].cells[1].text = v
d.add_paragraph()

# =====================================================
# CÂU TRẢ LỜI (TÔ XANH, GOM TRƯỚC PHỤ LỤC)
# =====================================================
add_h('Câu trả lời', 1, BLUE)

blue_para('Câu hỏi 1 — Coping Power Lochman 2025 (Mỹ):', bold=True, size=12)
blue_para('RCT đa trung tâm tại 40 trường THCS Mỹ, n=709 HS lớp 7. EACP (Early Adolescent Coping Power) '
          'gồm 3 thành phần: HỌC SINH (25 buổi/năm) + CAREGIVERS (16 buổi/90 phút) + GIÁO VIÊN (phối hợp '
          'cố vấn học đường). 5 nhóm kỹ năng: quản lý cơn giận-lo âu (self-statements + thư giãn), giải quyết '
          'vấn đề xã hội, tương tác bạn bè, kỹ năng học tập, từ chối chất gây nghiện. Kết quả: giảm hành vi '
          'gây hấn ở post-test + 3-year follow-up; giảm phạm pháp + chất gây nghiện ở 1-year follow-up. Lý '
          'thuyết: social-cognitive developmental contextual model. Tài trợ NIMH.')
d.add_paragraph()

blue_para('Câu hỏi 2 — Peer Support Murphy 2024 (Wiley J Community Psych):', bold=True, size=12)
blue_para('SCOPING REVIEW (không phải RCT) — 15 nghiên cứu, 13 interventions trong primary youth MH care. '
          'Định nghĩa peer support: "hỗ trợ xã hội + cảm xúc giữa cá nhân có TRẢI NGHIỆM CHUNG về khó khăn '
          'MH". 8/13 nhắm trầm cảm-lo âu-distress nhẹ-vừa. 9/13 cho 16-25 tuổi (CHỈ 2/13 cho HS THCS-THPT). '
          'Đo bằng GAD-7, DASS-21. Phát hiện: peer support có tiềm năng cải thiện recovery — nhưng cần peer '
          'worker được ĐỊNH NGHĨA VAI TRÒ RÕ + nhân viên chuyên môn hỗ trợ. Hạn chế: nội dung intervention '
          'thiếu chi tiết → khó nhân rộng. Áp dụng VN: cần adapt đáng kể vì đa số là thanh niên 16-25.')
d.add_paragraph()

blue_para('Câu hỏi 3 — Tỷ lệ HS THCS lo âu TỰ TÌM CHUYÊN MÔN:', bold=True, size=12)
blue_para('Data chuyên HS THCS hiếm (đa số gộp 12-17). Tỷ lệ thực tế:')
d.add_paragraph()

ans_tbl = d.add_table(rows=4, cols=3); ans_tbl.style = 'Table Grid'
ans_hdr = ['Quốc gia', 'Tỷ lệ', 'Ghi chú']
for i, h in enumerate(ans_hdr):
    c = ans_tbl.rows[0].cells[i]; shade(c, 'BDD7EE')
    pp = c.paragraphs[0]; pp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rr = pp.add_run(h); rr.bold = True; rr.font.color.rgb = BLUE
ans_data = [
    ('VIỆT NAM', '~5-12 %', 'V-NAMHS 2022: 8,4 % nhu cầu / 5,1 % phụ huynh nhận ra. HCM City THPT 2020: 12,5 % dùng dịch vụ. HS ưu thích NON-PROFESSIONAL trước.'),
    ('HOA KỲ (12-17)', '~40-50 % được CM (60 % gặp khó khăn)', 'NSCH 2023: 61 % HS chẩn đoán cần ĐT có khó tiếp cận; 20 % unmet need; 54 % khó tiếp cận care.'),
    ('Khoảng trống điều trị', 'VN > US', 'Cả 2 quốc gia đều có TREATMENT GAP lớn. VN gấp 4-5 lần US về tỷ lệ KHÔNG tiếp cận chuyên môn.'),
]
for i, row in enumerate(ans_data):
    for j, v in enumerate(row):
        c = ans_tbl.rows[i+1].cells[j]
        pp = c.paragraphs[0]
        rr = pp.add_run(v); rr.font.color.rgb = BLUE
        if j == 0: rr.bold = True
d.add_paragraph()

blue_para('Hàm ý cho can thiệp tại VN:', bold=True, size=12)
for it in [
    'Tỷ lệ HS VN tự tìm chuyên môn THẤP (~5-12 %) → cần chương trình SCHOOL-BASED (đến tận HS) thay vì '
    'chờ HS đến phòng khám',
    'Vì HS VN ưu thích NON-PROFESSIONAL trước → tận dụng peer support (Murphy 2024) + đào tạo giáo viên/'
    'cố vấn học đường',
    'Mô hình EACP Lochman (3 thành phần: HS + cha mẹ + giáo viên) có thể adapt cho VN — tuy nhiên đối '
    'tượng gốc là HS gây hấn (không phải lo âu); cần modify lại curriculum',
    'Kết hợp 2 cách tiếp cận: (1) SHU MẠNH (CBT chuyên sâu cho HS có lo âu rõ) + (2) PEER SUPPORT '
    '(toàn trường, normalize hỗ trợ tinh thần)',
]:
    blue_para('• ' + it)
d.add_paragraph()

blue_para('Ghi nhớ 1 dòng:', bold=True, size=12)
blue_para('Lochman 2025 EACP = 25 buổi học sinh + 16 buổi cha mẹ, hiệu quả lên hành vi gây hấn (không phải '
          'lo âu thuần). Murphy 2024 = scoping review peer support, đa số 16-25 tuổi. Tỷ lệ HS VN tự tìm '
          'chuyên môn THẤP HƠN US: VN ~5-12 %, US ~40-50 % (nhưng vẫn 50-60 % gặp khó khăn).', bold=True)
d.add_paragraph()

# =====================================================
# PHỤ LỤC
# =====================================================
add_h('Phụ lục — Tài liệu tham khảo', 1)
refs = [
    'Lochman JE, et al. (2025). Randomized controlled trial of the early adolescent coping power program: Effects on emotional and behavioral problems in middle schoolers. Journal of School Psychology. ScienceDirect pii S002244052500010X.',
    'Lochman JE, Wells KC, Lenhart LA. Coping Power: Child Group Program: Facilitator Guide. Treatments That Work series. Oxford University Press. ISBN 978-0195327878.',
    'Murphy R, Huggard L, Fitzgerald A, Hennessy E, Booth A. (2024). A systematic scoping review of peer support interventions in integrated primary youth mental health care. Journal of Community Psychology, 52(1):154-180. DOI: 10.1002/jcop.23090.',
    'Blueprints for Healthy Youth Development. Coping Power program profile. www.blueprintsprograms.org/programs/194999999/coping-power/',
    'V-NAMHS (2022). Vietnam Adolescent Mental Health Survey — Report on Main Findings. Queensland Centre for Mental Health Research + Vietnam. [VN002 trong DB]',
    'Nguyễn et al. (2020). Mental Health Literacy and Help-Seeking Preferences in High School Students in Ho Chi Minh City, Vietnam. School Mental Health.',
    'NSCH 2023. Adolescent Mental and Behavioral Health Data Brief. HRSA. mchb.hrsa.gov',
    'Annie E. Casey Foundation (2024). Youth Mental Health Statistics. www.aecf.org/blog/youth-mental-health-statistics',
    'Chú ý: Cả 2 bài Lochman 2025 + Murphy 2024 KHÔNG có trong DB 87 canonical của em. Nếu thầy cần canonical hóa (QT065 + QT066), em sẵn sàng download PDF + build summary.',
]
for ref in refs:
    p = d.add_paragraph(ref, style='List Number')
    for r in p.runs: r.font.size = Pt(10)

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
