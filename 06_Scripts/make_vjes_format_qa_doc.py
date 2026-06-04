# -*- coding: utf-8 -*-
"""Tạo file Word trả lời câu hỏi về quy định định dạng VJES cho thầy Minh Đức."""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/Trả lời quy định định dạng VJES.docx"
BLUE = RGBColor(0x00, 0x00, 0xCC)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(13)


def heading(t):
    p = doc.add_paragraph()
    r = p.add_run(t)
    r.bold = True
    r.font.size = Pt(13)
    return p


def body(t):
    p = doc.add_paragraph()
    p.add_run(t)
    return p


def answer(label, t):
    """Dòng kết luận — tô xanh."""
    p = doc.add_paragraph()
    r0 = p.add_run(label + " ")
    r0.bold = True
    r = p.add_run(t)
    r.bold = True
    r.font.color.rgb = BLUE
    return p


title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("TRẢ LỜI: QUY ĐỊNH ĐỊNH DẠNG VÀ TRÌNH BÀY CỦA TẠP CHÍ KHGDVN")
tr.bold = True
tr.font.size = Pt(14)

body("Kính gửi thầy Minh Đức,")
body("Dưới đây là trả lời bốn câu hỏi thầy nêu, căn cứ trực tiếp vào hai tài liệu "
     "chính thức của Tạp chí: “Tóm tắt Hướng dẫn dành cho tác giả 2025” và "
     "“Hướng dẫn cách trình bày, sắp xếp và trích dẫn tài liệu tham khảo theo "
     "quy định của APA”. Phần in xanh là kết luận thầy có thể dùng trực tiếp.")
doc.add_paragraph()

# 1
heading("1. Tạp chí quy định cỡ chữ (co chữ) bao nhiêu?")
body("Theo mục “Định dạng văn bản” trong Hướng dẫn dành cho tác giả, toàn bài dùng "
     "một phông chữ duy nhất là Times New Roman, với các cỡ chữ quy định cụ thể "
     "cho từng phần.")
answer("→ Cỡ chữ:",
       "Tên bài báo 12 pt in đậm; Tóm tắt và Từ khóa 11 pt (chữ thường, chỉ in "
       "đậm hai chữ “Tóm tắt” và “Từ khóa”); toàn bộ nội dung chính của bài 12 pt. "
       "Giãn dòng: Before 3 pt, After 3 pt, Line spacing Multiple 1,1. Độ dài "
       "bài 5.000–7.000 từ (khoảng 8–12 trang).")
doc.add_paragraph()

# 2
heading("2. Tài liệu tham khảo có dùng cỡ chữ nhỏ hơn không?")
body("Hướng dẫn chỉ quy định cỡ 11 pt cho riêng hai phần Tóm tắt và Từ khóa. Tất "
     "cả các phần còn lại của bài — bao gồm cả mục Tài liệu tham khảo — đều dùng "
     "cỡ 12 pt như nội dung chính. Tạp chí không có quy định thu nhỏ phần Tài "
     "liệu tham khảo.")
answer("→ Trả lời:",
       "Không. Mục Tài liệu tham khảo vẫn để cỡ 12 pt, không thu nhỏ. Chỉ Tóm "
       "tắt và Từ khóa mới dùng cỡ 11 pt.")
doc.add_paragraph()

# 3
heading("3. Cô Hằng có phải nộp phiếu khảo sát không, và có cần bảo mật không?")
body("Hướng dẫn nộp bài của Tạp chí không có quy định bắt buộc nộp kèm phiếu khảo "
     "sát; hồ sơ chỉ yêu cầu bản thảo bài báo và các tệp nguồn của bảng biểu, "
     "hình ảnh. Tuy nhiên, do biên tập viên đã nêu thắc mắc về công cụ đo lo âu "
     "(bản thang DSM-5 gốc gồm 10 mục, thang điểm 0–4, trong khi bài mô tả thang "
     "5 mục, thang điểm 1–4), việc nộp kèm phiếu khảo sát để biên tập đối chiếu "
     "là nên làm — giúp xác nhận đúng công cụ đã sử dụng.")
body("Về bảo mật: phiếu khảo sát ở dạng bản trống (bản công cụ đo) không chứa "
     "thông tin cá nhân của học sinh nên không phát sinh vấn đề bảo mật. Chỉ "
     "những phiếu đã điền (có dữ liệu danh tính học sinh) mới cần ẩn danh.")
answer("→ Đề xuất:",
       "Cô Hằng nên gửi kèm phiếu khảo sát — đúng bản đã dùng để thu dữ liệu 433 "
       "học sinh, khớp với mô tả trong bài. Bản phiếu trống không chứa dữ liệu cá "
       "nhân nên không cần bảo mật.")
doc.add_paragraph()

# 4
heading("4. Cụm “và cộng sự” có được viết tắt thành “cs.” không?")
body("Được. Hướng dẫn sắp xếp và trích dẫn tài liệu tham khảo của Tạp chí (các "
     "mục 4.7, 4.8 và Bảng tổng hợp 4.10) cho phép và chính Tạp chí cũng sử dụng "
     "cách viết tắt này, nhưng phân biệt theo vị trí trích dẫn:")
body("– Trích dẫn tường thuật (tên tác giả nằm trong câu, ngoài ngoặc đơn): viết "
     "đầy đủ “và cộng sự”. Ví dụ: Nguyễn Văn A và cộng sự (2026) cho thấy…")
body("– Trích dẫn trong ngoặc đơn: dùng viết tắt. Ví dụ: (Nguyễn Văn A & cs, "
     "2026).")
body("Lưu ý: Tạp chí viết là “cs” (KHÔNG có dấu chấm), đi kèm dấu “&” và chỉ dùng "
     "bên trong ngoặc đơn — không viết “cs.”. Chỉ dùng “cs” khi có từ hai tác giả "
     "trở lên bị lược.")
answer("→ Trả lời:",
       "Có, nhưng chỉ dùng trong ngoặc đơn và viết là “& cs” (không có dấu chấm); "
       "còn khi tên tác giả nằm trong câu thì phải viết đầy đủ “và cộng sự”.")
doc.add_paragraph()

body("Trân trọng.")

cp = doc.core_properties
cp.author = ""
cp.last_modified_by = ""
doc.save(OUT)
print("Đã lưu:", OUT)
