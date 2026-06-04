# -*- coding: utf-8 -*-
"""Bài 5: Alharbi et al. 2019 — theo phong cách CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 5 =====
bl('Tóm tắt bài 5', bold=True)

# BƯỚC 1: ĐỊNH DANH
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình nghiên cứu \u00ab Trầm cảm và lo âu ở học sinh trung học phổ thông tại vùng Qassim \u00bb (Depression and Anxiety among High School Student at Qassim Region), do Reem Alharbi, Khalid Alsuhaibani, Abdullah Almarshad và Abdulhameed Alyahya (2019), thuộc Trường Y khoa và Khoa Tâm thần, Đại học Qassim, Ả Rập Saudi, khảo sát trên mẫu 1.245 học sinh trung học phổ thông từ 13 đến 19 tuổi tại vùng al-Qassim, Vương quốc Ả Rập Saudi, xuất bản trên Journal of Family Medicine and Primary Care, tập 8, số 2, trang 504\u2013510.')

# BƯỚC 2: TỔNG QUAN PP + KT1
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng Bảng câu hỏi Sức khỏe Bệnh nhân 9 mục (PHQ-9 \u2014 Patient Health Questionnaire-9) để đánh giá trầm cảm và Thang đo Rối loạn Lo âu Tổng quát 7 mục (GAD-7 \u2014 Generalized Anxiety Disorder 7-item Scale) để đánh giá lo âu. Nói cách khác, công trình này sử dụng phương pháp sàng lọc cắt ngang bằng hai công cụ tiêu chuẩn quốc tế đã được xác thực.', bold=True)

# BƯỚC 3: ĐỊNH NGHĨA (KT5)
bl('PHQ-9 là thang đo tự báo cáo 9 mục dựa trên chín tiêu chí chẩn đoán trầm cảm chủ yếu của DSM-IV (Kroenke và cộng sự, 2001). Thang điểm từ 0 đến 27, phân loại thành năm mức: không có (0\u20134), nhẹ (5\u20139), trung bình (10\u201314), nặng vừa (15\u201319) và nặng (20\u201327).')
bl('GAD-7 là thang đo tự báo cáo 7 mục do Spitzer và cộng sự (2006) phát triển. Thang điểm từ 0 đến 21, với các ngưỡng cắt 5, 10 và 15 tương ứng với lo âu nhẹ, trung bình và nặng.')

# BƯỚC 4: BIỆN MINH (KT4)
bl('Tổng quan tài liệu của chúng tôi cho thấy không có nghiên cứu nào đo lường trầm cảm và lo âu ở tuổi vị thành niên tại vùng Qassim sử dụng PHQ-9 và GAD-7 \u2014 hầu hết các nghiên cứu tại Ả Rập Saudi trước đó sử dụng DASS (Thang đo Trầm cảm Lo âu Căng thẳng). Đây là nghiên cứu tiên phong sử dụng bộ công cụ này tại khu vực này.', bold=True)

# BƯỚC 5: LIỆT KÊ DỮ LIỆU
bl('Dữ liệu nhân khẩu học-xã hội bao gồm nhóm tuổi (13\u201314, 15\u201316, 17\u201318, ≥19), giới tính, nơi cư trú (9 khu vực), loại hệ thống giáo dục (phổ thông, hệ thống giáo trình, trường Quran), trình độ học vấn (năm 1\u20133) và tình trạng hôn nhân.')

# BƯỚC 6: MÔ TẢ QUY TRÌNH
bl('Nghiên cứu cắt ngang được thực hiện từ tháng 2/2018 đến tháng 5/2018 tại các trường trung học phổ thông vùng Qassim. Tiêu chí loại trừ: học sinh có khuyết tật trí tuệ. Dữ liệu được phân tích bằng SPSS với kiểm định chi bình phương, P ≤ 0,05 tại khoảng tin cậy 95%. Nghiên cứu được phê duyệt bởi Ủy ban Đạo đức Khu vực tại Đại học Công chúa Noura.')

# BƯỚC 9: ĐỊNH LƯỢNG — BẢNG
rb('Kết quả nghiên cứu định lượng')

bl('Phân bố trầm cảm theo PHQ-9 (N = 1.245):', bold=True)
tbl(
    ['Mức độ trầm cảm', 'n', 'Tỷ lệ'],
    [
        ['Không có (0\u20134)', '325', '26,0%'],
        ['Nhẹ (5\u20139)', '423', '34,0%'],
        ['Trung bình (10\u201314)', '306', '24,6%'],
        ['Nặng vừa (15\u201319)', '129', '10,4%'],
        ['Nặng (20\u201327)', '62', '5,0%'],
    ],
    widths=[5.0, 2.5, 2.5]
)
doc.add_paragraph()

bl('Phân bố lo âu theo GAD-7 (N = 1.245):', bold=True)
tbl(
    ['Mức độ lo âu', 'n', 'Tỷ lệ'],
    [
        ['Không có (0\u20134)', '455', '36,5%'],
        ['Nhẹ (5\u20139)', '425', '34,1%'],
        ['Trung bình (10\u201314)', '243', '19,5%'],
        ['Nặng (15\u201321)', '122', '9,8%'],
    ],
    widths=[5.0, 2.5, 2.5]
)
doc.add_paragraph()

bl('Mối liên quan giữa lo âu và các đặc điểm nhân khẩu học-xã hội:', bold=True)
tbl(
    ['Đặc điểm', 'P', 'Ý nghĩa'],
    [
        ['Giới tính (nữ > nam)', '<0,001', 'Có ý nghĩa'],
        ['Nơi cư trú', '<0,001', 'Có ý nghĩa'],
        ['Loại hệ thống giáo dục', '0,022', 'Có ý nghĩa'],
        ['Tình trạng hôn nhân', '0,001', 'Có ý nghĩa'],
        ['Nhóm tuổi', '0,341', 'Không ý nghĩa'],
        ['Trình độ học vấn (năm)', '0,545', 'Không ý nghĩa'],
        ['Mức trầm cảm \u2194 lo âu', '<0,001', 'Tương quan mạnh'],
    ],
    widths=[5.0, 2.5, 3.5]
)
doc.add_paragraph()

# BƯỚC 10: ĐỐI CHIẾU (KT3)
bh3('Khác biệt giới tính : Nữ giới luôn được báo cáo có tỷ lệ trầm cảm và lo âu cao hơn nam giới trong y văn quốc tế (Salk và cộng sự, 2017). Thực tế, tại Ả Rập Saudi, các nghiên cứu trước đó cũng xác nhận xu hướng này (Al-Gelban, 2007, 2009). Chúng tôi xác nhận trong mẫu 1.245 học sinh rằng nữ chiếm 71,3% nhóm lo âu nặng so với nam 28,7%, và 74,2% nhóm trầm cảm nặng so với nam 25,8% (P < 0,001 cho cả hai).')

# BƯỚC 11+12: NHẬN XÉT
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về tỷ lệ cao và đồng mắc trầm cảm-lo âu')
bl('*Tỷ lệ rất cao.* Nổi bật nhất là 74% học sinh có triệu chứng trầm cảm ở một mức nào đó (từ nhẹ đến nặng) và 63,5% có triệu chứng lo âu. Đây là tỷ lệ cao bất thường khi so với các nghiên cứu quốc tế sử dụng cùng công cụ: Chen và cộng sự (2023) báo cáo 23% trầm cảm và 13,9% lo âu tại Trung Quốc. Sự khác biệt này gợi ý rằng hoặc ngưỡng cắt (cut-off) của PHQ-9/GAD-7 cần được điều chỉnh cho bối cảnh văn hóa Ả Rập Saudi, hoặc tỷ lệ thực sự rất cao tại khu vực này.')
bl('*Đồng mắc trầm cảm và lo âu.* Mức độ trầm cảm và lo âu có tương quan mạnh và có ý nghĩa thống kê (P < 0,001). Trong nhóm trầm cảm nặng, 74,2% cũng có lo âu nặng. Phát hiện này phù hợp với y văn quốc tế về tỷ lệ đồng mắc cao giữa hai rối loạn ở thanh thiếu niên.')
bl('*Yếu tố nơi cư trú.* Nơi cư trú có mối liên quan có ý nghĩa thống kê với cả trầm cảm và lo âu (P < 0,001), với Buraidah (thành phố lớn nhất vùng) có tỷ lệ cao nhất. Điều này gợi ý ảnh hưởng của đô thị hóa đến sức khỏe tâm thần thanh thiếu niên, tương tự phát hiện của Xu và cộng sự (2021) về sự khác biệt thành thị-nông thôn tại Trung Quốc.')

# BƯỚC 13: KẾT LUẬN (Kiến trúc B)
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy tỷ lệ triệu chứng trầm cảm 74% và lo âu 63,5% ở 1.245 học sinh trung học phổ thông tại vùng Qassim \u2014 trong đó 15,4% trầm cảm nặng vừa đến nặng và 9,8% lo âu nặng, gợi ý rằng sức khỏe tâm thần thanh thiếu niên tại Ả Rập Saudi cần được chú ý khẩn cấp. Sự khác biệt giới tính rõ rệt (nữ > nam, P < 0,001) và tương quan mạnh giữa trầm cảm và lo âu nhấn mạnh nhu cầu sàng lọc cả hai tình trạng đồng thời. Trong bối cảnh văn hóa bảo thủ như Ả Rập Saudi, nhu cầu nâng cao nhận thức về sức khỏe tâm thần cho phụ huynh và giáo viên, cùng với phối hợp giữa Bộ Y tế và Bộ Giáo dục, là thiết yếu để giảm gánh nặng rối loạn tâm thần ở thanh thiếu niên.', bold=True)

# PHẢN BIỆN
rh2('Phản biện')
bl('Tỷ lệ 74% trầm cảm và 63,5% lo âu (bất kỳ mức độ nào) là rất cao so với y văn quốc tế \u2014 khi bao gồm cả mức "nhẹ" (PHQ-9: 5\u20139, GAD-7: 5\u20139), hầu hết dân số sẽ được phân loại là có triệu chứng, đặt câu hỏi về tính phân biệt lâm sàng (clinical discriminant validity) của ngưỡng cắt thấp. Nếu loại bỏ mức nhẹ, tỷ lệ trầm cảm trung bình\u2013nặng là 40% và lo âu trung bình\u2013nặng là 29,3% \u2014 vẫn cao nhưng hợp lý hơn. Chỉ sử dụng chi-square (đơn biến) mà không dùng hồi quy logistic đa biến. Phương pháp chọn mẫu không rõ ràng: bài báo chỉ nêu "dân số nghiên cứu bao gồm hầu hết các trường" mà không mô tả quy trình lấy mẫu. Tạp chí JFMPC (PubMed-indexed, PMID: 30984663) có chất lượng xuất bản tốt hơn nhiều so với tạp chí không chỉ mục, nhưng impact factor còn khiêm tốn.')

# HƯỚNG NC
rh2('Hướng nghiên cứu tiếp theo')
bl('Tái lặp nghiên cứu với phân tích hồi quy đa biến để xác định yếu tố nguy cơ độc lập. So sánh ngưỡng cắt PHQ-9/GAD-7 chuẩn quốc tế với ngưỡng được xác thực cho bối cảnh Ả Rập Saudi. Nghiên cứu yếu tố văn hóa đặc thù (áp lực gia đình, kỳ vọng xã hội) ảnh hưởng đến tỷ lệ cao bất thường. Nghiên cứu can thiệp phối hợp giữa nhà trường, gia đình và cơ sở y tế tại vùng Qassim.')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50 Trung bình-Khá. Cỡ mẫu lớn (1.245), sử dụng PHQ-9 và GAD-7 đã xác thực, PubMed-indexed, nhưng thiếu phân tích đa biến và tỷ lệ cao bất thường cần giải thích thêm.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('05_Alharbi_et_al_2019.docx')
print('05_Alharbi_et_al_2019.docx OK')
