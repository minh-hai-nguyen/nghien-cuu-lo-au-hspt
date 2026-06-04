# -*- coding: utf-8 -*-
"""Kiem tra noi dung tung dong - tim cac van de tiem an:
- Doan trong giua noi dung (khong mong doi)
- Heading nam giua body (tach saI)
- Text gan duplicate (co the la noi dung lap)
- Bullet list khong nhat quan
- Empty headings
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
d = Document(FILE)

print("="*80)
print("KIEM TRA NOI DUNG TUNG DONG")
print("="*80)

# ============================================================
# 1. Empty paragraphs WITH Heading style (orphan headings)
# ============================================================
print("\n[1] HEADING NHUNG TEXT TRONG (orphan heading)")
orphan_headings = []
for i, p in enumerate(d.paragraphs):
    text = p.text.strip()
    style = p.style.name if p.style else 'Normal'
    if ('Heading' in style or 'Tiêu đề' in style) and not text:
        orphan_headings.append((i, style))
print(f"  Tim thay: {len(orphan_headings)}")
for i, style in orphan_headings[:5]:
    print(f"    para {i} [{style}]: TRONG")

# ============================================================
# 2. Paragraphs that look like headings but Normal style
# ============================================================
print("\n[2] DOAN GIONG HEADING NHUNG STYLE NORMAL")
looks_like_heading = []
def in_table(p):
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag in (qn('w:tbl'), qn('w:tc')):
            return True
        parent = parent.getparent()
    return False

for i, p in enumerate(d.paragraphs):
    if i < 100: continue  # skip cover
    if in_table(p): continue
    text = p.text.strip()
    if not text or len(text) > 200: continue
    style = p.style.name if p.style else ''
    if 'Heading' in style or 'Tiêu đề' in style: continue
    # Check if looks heading
    upper = text.upper()
    # Pattern: numbered + short, OR all caps + short
    if re.match(r'^\d+\.\d+', text) and len(text) < 100:
        # Check bold-ness
        bold_runs = sum(1 for r in p.runs if r.text.strip() and r.bold)
        total_runs = sum(1 for r in p.runs if r.text.strip())
        if total_runs > 0 and bold_runs / total_runs > 0.5:
            looks_like_heading.append((i, text, style, 'numbered+bold'))
    elif upper == text and 3 < len(text) < 80 and not text.startswith(('Bảng', 'Hình')):
        # All caps short
        looks_like_heading.append((i, text, style, 'all-caps-short'))

print(f"  Tim thay: {len(looks_like_heading)}")
for i, text, style, reason in looks_like_heading[:15]:
    print(f"    para {i} [{style}] ({reason}): {text[:80]}")

# ============================================================
# 3. Empty/blank paragraphs in middle of content (between heading and next heading)
# ============================================================
print("\n[3] DOAN TRONG GIUA NOI DUNG (>= 3 dong trong lien tiep)")
clusters = []
empty_start = None
empty_count = 0
for i, p in enumerate(d.paragraphs):
    if not p.text.strip():
        if empty_start is None:
            empty_start = i
        empty_count += 1
    else:
        if empty_count >= 3:
            clusters.append((empty_start, i-1, empty_count))
        empty_start = None
        empty_count = 0
print(f"  Tim thay: {len(clusters)} cum >= 3 doan trong")
for start, end, n in clusters[:10]:
    # Show context: previous non-empty text
    prev_idx = start - 1
    while prev_idx > 0 and not d.paragraphs[prev_idx].text.strip():
        prev_idx -= 1
    prev = d.paragraphs[prev_idx].text.strip() if prev_idx >= 0 else '?'
    next_idx = end + 1
    nxt = d.paragraphs[next_idx].text.strip() if next_idx < len(d.paragraphs) else '?'
    print(f"    para {start}-{end} ({n} trong) - sau: {prev[:40]!r} | truoc: {nxt[:40]!r}")

# ============================================================
# 4. Headings out of order (1.2 before 1.1 etc.)
# ============================================================
print("\n[4] HEADING SAI THU TU (xuat hien khong dung trinh tu)")
prev_num = None
prev_idx = None
out_of_order = []
for i, p in enumerate(d.paragraphs):
    style = p.style.name if p.style else ''
    if 'Heading' not in style and 'Tiêu đề' not in style: continue
    text = p.text.strip()
    m = re.match(r'^(\d+(?:\.\d+)*)', text)
    if not m: continue
    num = m.group(1)
    parts = [int(x) for x in num.split('.')]
    if prev_num is not None:
        # Check if current is "logically after" previous
        # E.g., 1.1.5 after 1.1.4 OK, 1.1.4 after 1.1.5 NOT OK (within same parent)
        prev_parts = [int(x) for x in prev_num.split('.')]
        # Compare lexicographic
        if parts < prev_parts and len(parts) >= len(prev_parts) - 1:
            # Going backwards in numbering - suspicious
            common_len = min(len(parts), len(prev_parts))
            if parts[:common_len-1] == prev_parts[:common_len-1]:
                # Same parent → ordering violation
                out_of_order.append((i, num, prev_idx, prev_num, text[:60]))
    prev_num = num
    prev_idx = i
print(f"  Tim thay: {len(out_of_order)}")
for i, num, pi, pn, txt in out_of_order[:5]:
    print(f"    para {i} ({num}) sau para {pi} ({pn}): {txt}")

# ============================================================
# 5. Very short heading text (could be incomplete)
# ============================================================
print("\n[5] HEADING QUA NGAN (co the chua hoan thanh)")
short_headings = []
for i, p in enumerate(d.paragraphs):
    style = p.style.name if p.style else ''
    if 'Heading' not in style and 'Tiêu đề' not in style: continue
    text = p.text.strip()
    if 1 < len(text) < 8 and not re.match(r'^CHƯƠNG\s+\d+$', text):
        short_headings.append((i, text, style))
print(f"  Tim thay: {len(short_headings)}")
for i, txt, style in short_headings[:10]:
    print(f"    para {i} [{style}]: {txt!r}")

# ============================================================
# Stats
# ============================================================
print()
print(f"  Total paragraphs: {len(d.paragraphs)}")
print(f"  Total tables: {len(d.tables)}")
print(f"  File size: {os.path.getsize(FILE)//1024} KB")
