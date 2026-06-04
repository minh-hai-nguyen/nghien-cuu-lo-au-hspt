# -*- coding: utf-8 -*-
"""
TECHNICAL METHODS APPENDIX — Noi bo
Tai lieu rieng ghi lai phan ky thuat cua he thong QA cho du an Lo au:
- He thong dat ten canonical VN001/QT001
- Knowledge Graph v1/v2 (NetworkX)
- RAG v6/v7 (ChromaDB + BGE-M3)
- QA multi-layer pipeline
- Version history v1->v4b->v5
- Orphan facts traceability
Tep nay KHONG gui sep — la tai lieu noi bo "how we built it".
"""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DST = os.path.join(ROOT, '01_Bao-cao',
                   'Bao cao Can thiep - Technical_Methods - 12042026.docx')
CHARTS = os.path.join(os.path.dirname(__file__), 'charts')
KG_DIR = os.path.join(os.path.dirname(__file__), 'kg_data')
PAGE_W = 16.0

doc = Document()
st = doc.styles['Normal']
st.font.name = 'Times New Roman'; st.font.size = Pt(12)
st.paragraph_format.space_after = Pt(6); st.paragraph_format.line_spacing = 1.5
for sec in doc.sections:
    sec.top_margin = Cm(2.5); sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3); sec.right_margin = Cm(2)

# ==================== helpers ====================
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcW.findall(qn('w:tcW')): tcW.remove(old)
    tcW.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def H(text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)

def P(text, bold=False, italic=False, size=12, align='justify'):
    p = doc.add_paragraph()
    if align == 'justify':
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    elif align == 'center':
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic

def table(headers, rows, widths):
    assert sum(widths) <= PAGE_W + 0.05
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    t.autofit = False
    set_grid(t, widths)
    for row in t.rows:
        for ci in range(len(headers)):
            colw(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name='Times New Roman'; r.font.size=Pt(10)
    return t

def img(filename, width_cm=15.5, caption=None):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    path = os.path.join(CHARTS, filename)
    if not os.path.exists(path):
        path = os.path.join(KG_DIR, filename)
    if os.path.exists(path):
        p.add_run().add_picture(path, width=Cm(width_cm))
    if caption:
        cp = doc.add_paragraph()
        cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = cp.add_run(caption)
        r.font.name='Times New Roman'; r.font.size=Pt(10); r.italic=True

# ============================================================
# COVER
# ============================================================
H('PHỤ LỤC KỸ THUẬT — HỆ THỐNG QA DỰ ÁN LO ÂU VTN', level=1)
P('Tài liệu nội bộ — ghi lại kiến trúc kỹ thuật của pipeline quản lý 68 bài nghiên cứu, '
  'Knowledge Graph, Retrieval-Augmented Generation (RAG), và hệ thống kiểm tra chất '
  'lượng nhiều lớp (multi-layer QA) sử dụng trong quá trình biên soạn báo cáo can thiệp '
  'lo âu vị thành niên (v1 → v5).',
  italic=True)
P('Cập nhật: 12/04/2026 | Tác giả: Nhóm kỹ thuật dự án Lo âu', italic=True, align='center')
P('LƯU Ý: Tài liệu này chỉ dành cho tham khảo kỹ thuật nội bộ, không thuộc nội dung báo '
  'cáo nghiên cứu chính (v5). Mục đích là để đồng nghiệp tiếp quản dự án có thể hiểu '
  'được cách hệ thống được xây dựng, reproduce, và mở rộng.',
  italic=True)

doc.add_page_break()

# ============================================================
# 1. TONG QUAN KIEN TRUC
# ============================================================
H('1. Tổng quan kiến trúc hệ thống', level=1)

P('Pipeline QA dự án Lo âu gồm 4 thành phần chính hoạt động song song:')

table(
    ['Thành phần', 'Vai trò', 'Công nghệ', 'Trạng thái'],
    [
        ['Canonical Naming',
         'Chuẩn hoá tên file toàn bộ tài liệu',
         'Regex + JSON mapping',
         'v1 — 11/04/2026'],
        ['Knowledge Graph',
         'Trích xuất entity/relation, phát hiện orphan facts',
         'NetworkX + RDFLib',
         'v1 (11/04) → v2 (12/04)'],
        ['Vector RAG',
         'Retrieval ngữ nghĩa đa ngôn ngữ',
         'ChromaDB + sentence-transformers',
         'v6 MiniLM → v7 BGE-M3'],
        ['Multi-layer QA',
         'Kiểm tra đồng thời 3 lớp: regex / RAG / KG',
         'Python pipeline',
         'qa_advanced_v3.py'],
    ],
    widths=[3.5, 5.5, 4.5, 2.0])

P('Nguyên tắc thiết kế: mỗi lớp QA độc lập — lỗi bỏ sót ở lớp này có thể được lớp khác '
  'bắt. Kết quả: báo cáo v3 → v4 → v4b phát hiện 33 orphan facts + 1 ghost claim + 16 '
  'viết tắt chưa định nghĩa.')

# ============================================================
# 2. CANONICAL NAMING
# ============================================================
H('2. Hệ thống đặt tên canonical (v1 — 11/04/2026)', level=1)

P('Vấn đề: Trước 11/04/2026, các file được đặt tên không nhất quán (vd. "Hoa2024.pdf", '
  '"NC_Hoa_HaNoi_2024.docx", "Hoa_Frontiers.pdf") — khiến cross-reference giữa PDF gốc, '
  'bản dịch và tóm tắt rất khó. Nhiều file trùng lặp tồn tại dưới tên khác nhau.')

P('Giải pháp: Hệ thống canonical với cấu trúc:', bold=True)
P('  {VN|QT}{3-digit}_{TacGia}_{ChuDe_hoac_Diadiem}_{NamXB}.{pdf|docx}')
P('Ví dụ:')
P('  • VN001_Hoa_2024_Frontiers_HaNoi.docx')
P('  • QT043_Bress_JAMA_Maya_App_2024.pdf')
P('  • VN030_Tran_HappyHouse_Cambridge_2023.docx')

P('Quy tắc cụ thể:')
P('  1. Prefix VN cho 30 bài Việt Nam, QT cho 51 bài quốc tế')
P('  2. Padding 3 chữ số (hỗ trợ mở rộng đến 999 bài/loại)')
P('  3. Descriptor viết liền bằng dấu gạch dưới, không dùng dấu tiếng Việt')
P('  4. Tóm tắt: Tom-tat-tung-bai/{ID}_{Descriptor}.docx')
P('  5. Bản dịch: 03_Ban-dich/{ID}_{Descriptor}.docx')
P('  6. PDF gốc: 02_Papers-goc/{Folder}/{ID}_{Descriptor}.pdf')

P('Thư mục PDF được phân loại theo địa lý:')
P('  • Viet-Nam/')
P('  • The-gioi_Au-My-Uc/')
P('  • The-gioi_Khac/')
P('  • Dong-Nam-A/')
P('  • 11-bai-ban-dau-va-mo-rong/')

P('Artifacts:', bold=True)
P('  • 02_Papers-goc/canonical_index.json — bảng tra cứu {canonical_id → metadata}')
P('  • 02_Papers-goc/canonical_mapping.json — ánh xạ tên cũ → tên mới (backup)')
P('  • 02_Papers-goc/MASTER_INDEX.md — danh sách đầy đủ với checkbox tóm tắt/dịch/PDF')
P('  • 06_Scripts/build_master_index.py — script regenerate MASTER_INDEX từ disk scan')

P('Kết quả sau khi áp dụng:')
P('  • 68 canonical IDs (20 VN + 48 QT sau rebuild ngày 12/04)')
P('  • 55 PDF gốc / 61 bản dịch / 68 tóm tắt')
P('  • 5 bài QT còn paywall (QT037, QT048–QT051) — đã document rõ trong MASTER_INDEX')

# ============================================================
# 3. KNOWLEDGE GRAPH
# ============================================================
doc.add_page_break()
H('3. Knowledge Graph v1 và v2', level=1)

H('3.1. Động lực xây dựng KG', level=2)
P('Hạn chế của regex: phụ thuộc format cụ thể — "5,730" vs "5.730" vs "5730" cần 3 '
  'pattern khác nhau. QA v3 gặp 6 false positive loại này.')
P('Hạn chế của RAG text-based: không bắt được RELATION giữa thực thể. Truy vấn "Papers '
  'using CBT to measure Social Anxiety in Vietnam" có thể retrieve chunks chứa CBT + '
  'Social Anxiety + Vietnam nhưng không đảm bảo thuộc cùng một bài.')
P('Knowledge Graph giải quyết cả hai vấn đề: entity + relation định nghĩa rõ ràng trong '
  'schema, queries phức tạp qua graph traversal.')

H('3.2. Schema', level=2)
P('KG sử dụng 9 loại nodes và 9 loại edges:')

table(
    ['Node type', 'Mô tả', 'Ví dụ'],
    [
        ['Paper', 'Bài nghiên cứu', 'VN030 (Tran Happy House)'],
        ['Author', 'Tác giả', 'Tran Thach Duc'],
        ['Journal', 'Tạp chí', 'Cambridge Global Mental Health'],
        ['Country', 'Quốc gia nghiên cứu', 'Vietnam'],
        ['Method', 'Phương pháp can thiệp', 'CBT, iCBT, RAP-A'],
        ['Outcome', 'Kết cục đo', 'Depression, Anxiety, SAD, Stress'],
        ['Scale', 'Công cụ đo', 'GAD-7, CESD-R, DASS-21'],
        ['EffectSize', 'Cỡ hiệu ứng', 'OR=0.56, d=0.11'],
        ['SampleSize', 'Cỡ mẫu', 'n=1084'],
    ],
    widths=[2.5, 5.5, 7.5])

P('Edge types:')
P('  AUTHORED_BY — Paper → Author')
P('  PUBLISHED_IN — Paper → Journal')
P('  CONDUCTED_IN — Paper → Country')
P('  USED_METHOD — Paper → Method')
P('  MEASURED — Paper → Outcome')
P('  USED_SCALE — Paper → Scale')
P('  REPORTED_ES — Paper → EffectSize')
P('  HAS_N — Paper → SampleSize')
P('  HAS_YEAR — Paper → Year')

H('3.3. KG v1 (11/04/2026)', level=2)
P('• File: 06_Scripts/kg_data/kg_v1.graphml')
P('• Script: 06_Scripts/kg_build_v1.py')
P('• Nodes: 218 (60 Paper, 62 EffectSize, 33 SampleSize, 19 Country, 13 Method, 13 '
  'Scale, 12 Year, 5 Outcome, 1 Journal)')
P('• Edges: 573')
P('• Validation rules R1–R8: 0 critical violations (06_Scripts/kg_validate_v1.py)')
P('• Build time: ~2 giây cho 60 papers')
P('• Vấn đề phát hiện: Country detection over-tag. Vietnam bị gán 50 papers thay vì ~21 '
  '— do match từ citation text thay vì field affiliation.')

H('3.4. KG v2 (12/04/2026) — Entity resolution fix', level=2)
P('• File: 06_Scripts/kg_data/kg_v2.graphml')
P('• Script: 06_Scripts/kg_build_v2.py')
P('• Nodes: 206 (giảm sau entity cleanup)')
P('• Edges: 492')
P('• Fix chính: Country detection chỉ scan 500 ký tự đầu + các field "Đơn vị", "Cơ '
  'quan", "Địa bàn", loại bỏ citation text. Vietnam papers giảm 50 → 21 (chính xác: '
  '19 VN001–029 + 2 outliers QT001/QT035 cần kiểm).')
P('• Limitation còn lại: First author extraction = 0 (format tóm tắt CTH v5 không có '
  'field "Tác giả" riêng). Cần regex từ full text cho KG v3.')

H('3.5. Cross-check với báo cáo v3 — phát hiện 33 orphan facts', level=2)
P('Script: 06_Scripts/kg_report_crosscheck.py')
P('Logic: với mỗi EffectSize node trong KG, search báo cáo v3 xem có số liệu gần tương '
  'đương (tolerance 0.01 cho OR) hay không. Nếu không có → orphan fact.')
P('Kết quả: 33 orphan facts + 1 ghost claim (báo cáo có số liệu nhưng KG không có → '
  'ghost). 18 orphan facts quan trọng nhất được đưa vào báo cáo v4b phần V.1.')

H('3.6. Visualization', level=2)
P('Script: 06_Scripts/kg_visualize.py')
P('• kg_interactive.html — pyvis network interactive (drag, zoom, filter)')
P('• kg_vietnam_subgraph.png — subgraph chỉ Vietnam papers')
P('• kg_method_outcome_heatmap.png — ma trận Method × Outcome (dùng trong báo cáo v4b/v5)')
P('Lỗi cp1252 encoding với pyvis đã được fix bằng cách manually gọi net.generate_html() '
  'rồi write với encoding="utf-8".')

# ============================================================
# 4. RAG
# ============================================================
doc.add_page_break()
H('4. Retrieval-Augmented Generation (RAG) — v6 và v7', level=1)

H('4.1. Lịch sử versioning', level=2)
table(
    ['Version', 'Model', 'Chunks', 'Source', 'Smoke test', 'Ngày'],
    [
        ['v3', 'MiniLM-L12', '~150', 'Báo cáo v3 only', '—', '10/04'],
        ['v4', 'MiniLM-L12', '206', 'Báo cáo v4', '4/5 (80%)', '11/04'],
        ['v5', 'MiniLM-L12', '206', 'Báo cáo v4 + canonical', '4/5 (80%)', '11/04'],
        ['v6', 'MiniLM-L12', '2.393', 'Report v4 + 68 sum + 61 trans', '9/10 (90%)', '12/04'],
        ['v7', 'BAAI/bge-m3', '2.409', 'Report v4 + 68 sum + 61 trans', '10/10 (100%)', '12/04'],
    ],
    widths=[1.8, 2.5, 1.8, 5.5, 2.6, 1.8])

H('4.2. Config v7 (production)', level=2)
P('• Collection: rag_v7_bgem3 trong rag_bao_cao_can_thiep/')
P('• Embedding model: BAAI/bge-m3 (1024 chiều, context window 8192 token)')
P('• Cache model: ~2.3 GB')
P('• Source documents: 130 file (1 report + 68 summaries + 61 translations)')
P('• Encoding time: 2013 giây (~1,2 chunks/giây trên CPU)')
P('• Script: 06_Scripts/rag_v7_bgem3.py')
P('• Manifest: 06_Scripts/rag_v7_manifest.json')

H('4.3. So sánh v7 vs v6', level=2)
P('Ưu điểm v7:', bold=True)
P('  • Query "Wen 2020 OR 11,58 rural China" PASS (v6 fail)')
P('  • Cross-lingual Việt → English retrieve tốt hơn')
P('  • Grain tinh hơn (1024 vs 384 chiều)')
P('  • State-of-the-art multilingual, hỗ trợ tốt tiếng Việt có dấu')

P('Nhược điểm v7:', bold=True)
P('  • Model lớn gấp 5× MiniLM → encode chậm ~3×')
P('  • Không deploy được trên Render.com free tier (512 MB RAM limit)')
P('  • Ổ đĩa chiếm ~15 MB cho collection (vs ~5 MB cho v6)')

H('4.4. Snippet sử dụng v7', level=2)
P('```python')
P('import chromadb')
P('from sentence_transformers import SentenceTransformer')
P('model = SentenceTransformer("BAAI/bge-m3")')
P('client = chromadb.PersistentClient(path="rag_bao_cao_can_thiep")')
P('col = client.get_collection("rag_v7_bgem3")')
P('q_emb = model.encode(["query tiếng Việt hoặc English"],')
P('                      normalize_embeddings=True).tolist()')
P('res = col.query(query_embeddings=q_emb, n_results=5,')
P('                include=["documents", "metadatas"])')
P('```')

# ============================================================
# 5. MULTI-LAYER QA
# ============================================================
doc.add_page_break()
H('5. Multi-layer QA Pipeline', level=1)

P('Script chính: 06_Scripts/qa_advanced_v3.py — 8 lớp kiểm tra, mỗi lớp bắt loại lỗi '
  'khác nhau.')

table(
    ['Lớp', 'Tên', 'Mô tả', 'Ví dụ lỗi bắt được'],
    [
        ['1', 'Consistency', 'Đối chiếu số liệu trong báo cáo với regex từ bản dịch',
         'OR=11.6 trong báo cáo khớp với VN001'],
        ['2', 'Plausibility', 'Check giá trị có hợp lý không',
         'OR > 100 hoặc tỷ lệ < 0% hoặc > 100%'],
        ['3', 'Citation', 'Mọi claim có citation hay không',
         'Câu có số liệu nhưng không có [VN001]'],
        ['4', 'Cross-ref summary', 'Số liệu trong báo cáo có match summary không',
         'Báo cáo viết 91.6% nhưng summary ghi 91.7%'],
        ['5', 'Interpretation', 'Diễn giải hướng đúng (OR > 1 risk, < 1 protective)',
         'Viết "yếu tố bảo vệ OR = 2.5" là sai'],
        ['6', 'Format', 'Format số Việt/Anh nhất quán',
         '5,730 vs 5.730 trong cùng một bảng'],
        ['7', 'Year', 'Năm publication vs data collection',
         'Brunborg Norway 2011/2025 — 2011 là data, 2025 là pub'],
        ['8', 'Abbreviation', 'Viết tắt định nghĩa lần đầu',
         'CBT, OR, AOR, SUCRA, MHST — 16 chưa define'],
    ],
    widths=[0.8, 2.7, 5.5, 7.0])

P('Kết quả QA v4 trên báo cáo v4: 36 issues phát hiện')
P('  • Critical: 0 real (10 false positive "n=1" do section numbering)')
P('  • Real issues đã fix: 2 (VN030 summary cập nhật, VN030 Happy House 12.881 char translation)')
P('  • 16 viết tắt chưa định nghĩa — cần glossary ở đầu báo cáo (pending)')
P('  • False positive: "Norway 2011/2025" — data year vs pub year (đúng, giữ)')

# ============================================================
# 6. VERSION HISTORY
# ============================================================
doc.add_page_break()
H('6. Lịch sử phiên bản báo cáo can thiệp', level=1)

table(
    ['Version', 'Ngày', 'Nội dung chính', 'Số paragraph', 'Số bảng', 'Char'],
    [
        ['v1', '07/04', 'Draft ban đầu, 21 papers', '~150', '5', '~20K'],
        ['v2', '10/04', '21 papers + phản biện đỏ', '~180', '6', '~25K'],
        ['v3', '11/04', '+ 9 papers mới (30 total), CTH v5', '~220', '8', '~32K'],
        ['v4', '12/04 sáng', '+ KG orphan facts, 5 insights compact', '254', '9', '~38K'],
        ['v4b', '12/04 trưa', 'v4 expanded (diễn giải sâu)', '355', '10', '52K'],
        ['v5', '12/04 chiều', 'Bản sạch gửi sếp (không có KG/RAG/AI)', '317', '9', '48K'],
    ],
    widths=[1.5, 2.0, 5.5, 2.3, 1.7, 3.0])

P('Lưu ý quan trọng về v5:', bold=True)
P('v5 là bản dành riêng cho cấp lãnh đạo — đã strip toàn bộ meta-information về tools, '
  'KG, RAG, AI, version comparison. Nội dung thuần tuý nghiên cứu. v4b vẫn giữ là bản '
  'nội bộ kỹ thuật đầy đủ với phần V.0 và V.5 về Knowledge Graph + RAG stack.')

P('Output paths:')
P('  • v3: 01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 11042026 v3.docx')
P('  • v4: 01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 12042026 v4.docx')
P('  • v4b: 01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 12042026 v4b.docx')
P('  • v5 (clean): 01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 12042026 v5.docx')
P('  • Technical (file này): 01_Bao-cao/Bao cao Can thiep - Technical_Methods - 12042026.docx')

# ============================================================
# 7. WORKFLOW HANG NGAY
# ============================================================
H('7. Workflow hằng ngày cho dự án (từ 12/04/2026)', level=1)

P('Khi thêm một paper mới vào hệ thống:')
P('  1. Đặt tên canonical {VN|QT}{3-digit}_{Descriptor}')
P('  2. Chuyển PDF vào 02_Papers-goc/{Folder}/')
P('  3. Dịch đầy đủ → 03_Ban-dich/{ID}_{Descriptor}.docx')
P('  4. Tóm tắt CTH v5 → Tom-tat-tung-bai/{ID}_{Descriptor}.docx')
P('  5. Chạy build_master_index.py → regenerate canonical_index.json + MASTER_INDEX.md')
P('  6. Chạy kg_build_v2.py → rebuild KG')
P('  7. Chạy rag_v7_bgem3.py → reindex RAG')
P('  8. Cập nhật báo cáo nếu cần')

P('Khi sửa báo cáo (bất kỳ version nào):')
P('  1. Chạy qa_advanced_v3.py (sửa path trong script) để check 8 lớp')
P('  2. Chạy kg_report_crosscheck.py để check orphan facts mới')
P('  3. Test smoke query trên RAG v7 với các claim chính')
P('  4. Release chỉ khi cả 3 lớp cho kết quả clean')

# ============================================================
# 8. BACKLOG KY THUAT
# ============================================================
H('8. Backlog kỹ thuật', level=1)

P('(a) KG v3 — First author extraction.', bold=True)
P('Hiện tại KG v2 không extract được first author do format CTH v5 không có field "Tác '
  'giả" riêng. Cần viết regex từ full text + author disambiguation (Brown 2024 BESST vs '
  'Brown & Carter 2025 Editorial).')

P('(b) Glossary 16 viết tắt.', bold=True)
P('QA v4 layer 8 phát hiện CBT, gCBT, DMHI, VRET, NMA, MA, SR, RCT, PE, OR, AOR, SUCRA, '
  'SMD, MHST, BESST, PLACES chưa định nghĩa lần đầu xuất hiện. Cần thêm glossary ở đầu '
  'báo cáo v5 (hoặc tạo phụ lục riêng).')

P('(c) 5 paywall PDFs.', bold=True)
P('QT037 Praptomojati (Wiley 403), QT048 Chen COVID Meta, QT049 Zhang Bayesian, QT050 '
  'Qiaochu Mobile CBT (Wiley 403), QT051 Menon LMIC. Cần truy cập qua thư viện đại học.')

P('(d) RAG v7 deployment.', bold=True)
P('BGE-M3 không fit Render.com free tier 512 MB. Hai lựa chọn: (a) upgrade paid tier; '
  '(b) dùng v6 MiniLM cho production, v7 cho local dev.')

P('(e) Integration với web assistant.', bold=True)
P('RAG v7 có thể gắn vào web trợ lý nghiên cứu (xem memory project_web_tro_ly_NC.md) '
  'làm backend retrieval. Cần API wrapper.')

# ============================================================
# APPENDIX: File inventory
# ============================================================
doc.add_page_break()
H('Phụ lục — File inventory', level=1)

P('Scripts (06_Scripts/):', bold=True)
P('  • build_master_index.py — regenerate MASTER_INDEX.md từ disk scan')
P('  • kg_build_v1.py, kg_build_v2.py — build Knowledge Graph')
P('  • kg_validate_v1.py — check rules R1-R8')
P('  • kg_report_crosscheck.py — detect orphan facts')
P('  • kg_visualize.py — pyvis HTML + PNG heatmap')
P('  • rag_v6_multisource.py — build RAG v6 (MiniLM)')
P('  • rag_v7_bgem3.py — build RAG v7 (BGE-M3)')
P('  • qa_advanced_v3.py — 8-layer QA pipeline')
P('  • bao_cao_v4b_expanded.py — sinh báo cáo v4b (nội bộ có KG/RAG)')
P('  • bao_cao_v5_clean.py — sinh báo cáo v5 (bản sạch cho sếp)')
P('  • bao_cao_technical_methods.py — sinh file này')
P('  • dich_VN030_HappyHouse.py — dịch full text VN030')
P('  • update_VN030_summary.py — update tóm tắt VN030')
P('  • tao_tom_tat_QT012_017.py — tạo 6 tóm tắt QT012–QT017')

P('Data (06_Scripts/kg_data/):', bold=True)
P('  • kg_v1.graphml, kg_v2.graphml — Knowledge Graph files')
P('  • kg_interactive.html — pyvis interactive')
P('  • kg_vietnam_subgraph.png — subgraph Vietnam')
P('  • kg_method_outcome_heatmap.png — Method × Outcome heatmap')

P('Indexes (02_Papers-goc/):', bold=True)
P('  • canonical_index.json — {id → metadata}')
P('  • canonical_mapping.json — {old_name → new_name}')
P('  • MASTER_INDEX.md — danh sách 68 bài với checkbox')

P('RAG collections (rag_bao_cao_can_thiep/):', bold=True)
P('  • rag_v6_multisource — ChromaDB collection, MiniLM-L12, 2.393 chunks')
P('  • rag_v7_bgem3 — ChromaDB collection, BGE-M3, 2.409 chunks')

P('Manifests (06_Scripts/):', bold=True)
P('  • rag_v6_manifest.json, rag_v7_manifest.json — metadata chunks')

P('---', align='center', italic=True)
P('Technical Methods Appendix — 12/04/2026 — nội bộ dự án Lo âu VTN',
  align='center', italic=True, size=10)

doc.save(DST)
d2 = Document(DST)
print(f'Saved: {os.path.basename(DST)}')
print(f'Total: {len(d2.paragraphs)} paragraphs, {len(d2.tables)} tables')
total_chars = sum(len(p.text) for p in d2.paragraphs)
print(f'Total chars: {total_chars:,}')
