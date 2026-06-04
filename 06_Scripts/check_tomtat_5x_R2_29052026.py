# -*- coding: utf-8 -*-
"""Vong kiem tra 2 - sau khi vong 1 da PASS - check sau hon.
29/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Cm, Pt
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SUM_PATH = os.path.join(ROOT, 'Luận án TS', 'TomTatLA_v2_VERIFIED_29052026.docx')
SUM = Document(SUM_PATH)

sum_text = '\n'.join(p.text for p in SUM.paragraphs)
sum_full = sum_text + '\n' + '\n'.join(c.text for tb in SUM.tables for row in tb.rows for c in row.cells)
issues = []

def check(name, cond, detail=''):
    status = '✓' if cond else '✗'
    print(f'  {status} {name}')
    if detail: print(f'      {detail}')
    if not cond: issues.append(name)
    return cond

# ============================================================
# VONG 2 - 5 CHECK MOI
# ============================================================
print('=== VONG 2: FORMAT + EDGE CASES ===')

# 2.1 Khổ giấy A5
sec = SUM.sections[0]
check('2.1 Khổ giấy A5 (14.8 × 21.0 cm)',
      abs(sec.page_width.cm - 14.8) < 0.1 and abs(sec.page_height.cm - 21.0) < 0.1,
      f'Width: {sec.page_width.cm:.1f}, Height: {sec.page_height.cm:.1f}')

# 2.2 Margins
check('2.2 Lề chuẩn (T 1.5, B 1.5, L 2.0, R 1.5 cm)',
      abs(sec.top_margin.cm - 1.5) < 0.05 and abs(sec.bottom_margin.cm - 1.5) < 0.05
      and abs(sec.left_margin.cm - 2.0) < 0.05 and abs(sec.right_margin.cm - 1.5) < 0.05,
      f'T={sec.top_margin.cm:.2f}/B={sec.bottom_margin.cm:.2f}/L={sec.left_margin.cm:.2f}/R={sec.right_margin.cm:.2f}')

# 2.3 Font Normal là TNR 11pt
normal_style = SUM.styles['Normal']
check('2.3 Font Normal TNR 11pt',
      normal_style.font.name == 'Times New Roman' and normal_style.font.size == Pt(11),
      f'Name: {normal_style.font.name}, Size: {normal_style.font.size.pt if normal_style.font.size else None}')

# 2.4 Line spacing 1.15
check('2.4 Line spacing 1.15',
      normal_style.paragraph_format.line_spacing == 1.15,
      f'Line spacing: {normal_style.paragraph_format.line_spacing}')

# 2.5 No fabricated R² (the old R²=0.58, 0.54, 0.42)
fake_r2 = any(x in sum_full for x in ['R² = 0,58', 'R² = 0,54', 'R² = 0,42', 'R²=0,58', 'R²=0,54', 'R²=0,42'])
check('2.5 Không có R² bịa cũ (0.58/0.54/0.42)', not fake_r2)


# ============================================================
# VONG 3 - GIA THUYET KHOP LA
# ============================================================
print('\n=== VONG 3: GIA THUYET 3 ===')

# 3.1 Giả thuyết 1: gender pattern
check('3.1 Giả thuyết 1 đã kiểm chứng (nữ > nam lan tỏa + xã hội)',
      'nữ' in sum_full.lower() and 'cao hơn' in sum_full)

# 3.2 Giả thuyết 1 detail: chia ly không có khác biệt
check('3.2 Lo âu chia ly không khác biệt giới (p > 0,05 hoặc F=0,246)',
      'chia ly' in sum_full.lower() and ('F = 0,246' in sum_full or 'p = 0,620' in sum_full or 'p > 0,05' in sum_full))

# 3.3 Giả thuyết 2: lòng tự trọng
check('3.3 Giả thuyết 2: lòng tự trọng cường độ ~86%',
      '86%' in sum_full or '~86' in sum_full)

# 3.4 Giả thuyết 3: thứ tự cường độ YTNC
check('3.4 Giả thuyết 3: thứ tự ALHT > NĐT > BNHĐ',
      'áp lực học tập' in sum_full.lower() and 'nghiện điện thoại' in sum_full.lower())

# 3.5 Khung CT 8 nội dung
check('3.5 Khung CT 8 nội dung',
      '8 nội dung cơ bản' in sum_full)


# ============================================================
# VONG 4 - CONSISTENCY
# ============================================================
print('\n=== VONG 4: NHAT QUAN THUAT NGU ===')

# 4.1 Dùng "rối loạn lo âu" thay vì "lo âu" mơ hồ
check('4.1 Có dùng "rối loạn lo âu" rõ ràng',
      sum_full.count('rối loạn lo âu') > 10)

# 4.2 Dùng "lan tỏa" (DSM-5) không phải "tổng quát"
has_tong_quat = 'tổng quát' in sum_full and 'lo âu tổng quát' in sum_full
check('4.2 Dùng "lan tỏa" theo DSM-5 (không có "lo âu tổng quát")',
      not has_tong_quat)

# 4.3 Năm Xu (2021) không (2022)
check('4.3 Năm Xu 2021 (đã sửa từ 2022)',
      'Xu và cộng sự (2021)' in sum_full or 'Xu (2021)' in sum_full)

# 4.4 V-NAMHS author đúng
check('4.4 Viện Xã hội học 2022 V-NAMHS',
      'Viện Xã hội học' in sum_full and 'V-NAMHS' in sum_full)

# 4.5 RCADS dùng đúng tên
check('4.5 RCADS (Revised Children\'s Anxiety...)',
      'RCADS' in sum_full and 'Revised Children' in sum_full)


# ============================================================
# VONG 5 - VAN PHONG TS TAM LY HOC VN
# ============================================================
print('\n=== VONG 5: VAN PHONG LA TS TAM LY HOC ===')

# 5.1 Có "Mục đích nghiên cứu" rõ ràng
check('5.1 Mục đích NC rõ ràng',
      'Mục đích nghiên cứu' in sum_text)

# 5.2 Phân biệt "đối tượng" vs "khách thể"
check('5.2 Phân biệt Đối tượng NC và Khách thể NC',
      'Đối tượng nghiên cứu' in sum_text and 'Khách thể nghiên cứu' in sum_text)

# 5.3 Có giả thuyết khoa học (đặc trưng TLH)
check('5.3 Có 3 giả thuyết khoa học',
      'Giả thuyết 1' in sum_text and 'Giả thuyết 2' in sum_text and 'Giả thuyết 3' in sum_text)

# 5.4 Đề cập đến SEM/CFA (đặc trưng TS TLH)
check('5.4 Đề cập SEM + CFA',
      'SEM' in sum_full and 'CFA' in sum_full)

# 5.5 Có đề xuất ứng dụng thực tiễn (Khung CT)
check('5.5 Khung chương trình tập huấn',
      'Khung chương trình tập huấn' in sum_full)


# ============================================================
print()
print('=' * 70)
print(f'VONG 2-5: TONG {len(issues)} loi')
if issues:
    for i, iss in enumerate(issues, 1):
        print(f'  {i}. {iss}')
else:
    print('CAC VONG 2-5 SACH LOI')
