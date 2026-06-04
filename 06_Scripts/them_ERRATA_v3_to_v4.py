# -*- coding: utf-8 -*-
"""Them ERRATA v3->v4 vao cuoi file v3."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
F = os.path.join(ROOT, 'bai-bao-Q1', 'ThamKhao_Titles_Q1Q3_AsiaChauPhi_TiengViet_v3_01062026.docx')

d = Document(F)
d.add_page_break()

def H1(t):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(10)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def H2(t):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(12); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

def P(t, italic=False, indent=False):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(t); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic


H1('ERRATA V3 → V4 — 25 SỬA THÊM (sau audit lần 3, 01/06/2026)')

P('Sau khi tạo v3, em đã thực hiện vòng audit thứ 3 theo yêu cầu "kiểm '
  'tra từng từ, từng câu một". Em search trực tiếp PubMed/Crossref/JCR '
  'cho TỪNG paper trong v3 và phát hiện thêm 25 vấn đề (9 lỗi nghiêm '
  'trọng + 16 chi tiết thiếu).', italic=True)

P('Tất cả các sửa này đã được tích hợp vào file v4 '
  '(ThamKhao_Titles_Q1Q3_AsiaChauPhi_TiengViet_v4_01062026.docx). Phần '
  'ERRATA này chỉ ghi nhận để minh bạch.', italic=True)


H2('9 LỖI NGHIÊM TRỌNG TRONG V3')

errors = [
    ('Q1#2 sai tác giả',
     'V3: "Liu Y, Xu L, Hagströmer M và cs."',
     'V4: Liu X, Yang F, Huang N, Zhang S, Guo J (tác giả đầu là Liu X, '
     'không phải Liu Y)'),
    ('Q1#4 sai tạp chí',
     'V3: "Frontiers in Psychiatry (Q1, IF 3.2)"',
     'V4: Frontiers in PSYCHOLOGY (Q2, IF 2.9) — Wang H, Qin Q 2025'),
    ('Q1#7 sai tạp chí + năm + IF',
     'V3: "Journal of Affective Disorders (Q1, IF 5.42) 2025"',
     'V4: Journal of ANXIETY Disorders 2026 — Kexin Zhang et al. — '
     'IF 4.5 (JCR 2024)'),
    ('Q1#10 sai tạp chí',
     'V3: "PMC (Frontiers Psychiatry)"',
     'V4: Frontiers in PSYCHOLOGY — Qianying Hu, Yingyan Zhong et al.'),
    ('Q1#17 sai tạp chí',
     'V3: "PMC (BMC Psychiatry) 2023"',
     'V4: Child Adolescent Psychiatry and Mental Health (Q1, IF 4.6) — '
     'Mbanuzuru V et al.'),
    ('Q3#3 sai tạp chí',
     'V3: "PMC (generic)"',
     'V4: Journal of Family Medicine and Primary Care — Sonam Kumari et al.'),
    ('Q3#5 sai publisher',
     'V3: "BMC 2024"',
     'V4: Psicologia: Reflexão e Crítica (Springer/Brazil, IF 2.0) — '
     'Ijaz S, Rohail I, Irfan S'),
    ('Q3#10 sai năm',
     'V3: "2014"',
     'V4: 2013 (Dec) — Dat Tan Nguyen, Christine Dedding, Tam Thi Pham, '
     'Pamela Wright, Joske Bunders — BMC Public Health'),
    ('IF General Psychiatry',
     'V3: "IF 7.0"',
     'V4: IF 6.8 (JCR 2024 chính thức)'),
]
for title, v3, v4 in errors:
    H2(title)
    P(v3, italic=False); P(v4, italic=True)


H2('16 CHỖ THIẾU CHI TIẾT (đã bổ sung vào v4)')

details = [
    'Q1#3 thêm tác giả: Jue Xu, Hanmin Duan, Kang Qin, Bing Liu',
    'Q1#8 thêm: Li L et al. — n=1.711 từ 30 RCT — BMC Psy article 809',
    'Q1#9 thêm: Thanh Long Nguyen, Victor Storm, Lam Toi Phung et al.',
    'Q1#11 thêm: JMIR Public Health and Surveillance (cụ thể) — Meng Wang '
    'et al. — n=22.868',
    'Q1#13 thêm tác giả đầu: Shraddha Bajaria',
    'Q1#14 thêm tác giả đầu: Astrid Jörns-Presentati, Ann-Kathrin Napp',
    'Q1#15 thêm tác giả đầu: Robera Demissie Berhanu — n=847',
    'Q3#1 thêm tác giả: Md Saiful Islam, Md Estiar Rahman et al. — n=563',
    'Q3#2 thêm tác giả: Afifa Anjum, Sahadat Hossain et al. — n=2.313',
    'Q3#3 thêm tác giả: Sonam Kumari, Arun Mahapatra et al. — n=679',
    'Q3#4 thêm tác giả: Roni Khatun, Salma Akter Urme, Md Syful Islam — n=400',
    'Q3#6 thêm journal cụ thể: Health Science Reports (Wiley) — '
    'Rasalingam G et al.',
    'Q3#7 thêm journal cụ thể: BMC Psychiatry — Zhangming Chen et al. — '
    'n=63.205',
    'Q3#8 thêm tác giả: Tingting Dong, Yumei Wang, Yuanjie Lin — n=2.716',
    'Q3#14 thêm journal cụ thể: Discover Mental Health — Alia Albinali '
    'et al. — n=750',
    'Q3#15 thêm journal cụ thể: BMC Pediatrics — Nabeel Al-Yateem et al. — '
    'n=968',
]
for det in details:
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5); p.paragraph_format.space_after = Pt(3)
    r = p.add_run('▸ ' + det); r.font.name = 'Times New Roman'; r.font.size = Pt(11)


H2('Bài học và quy tắc đã ghi vào memory')

P('Em đã viết memory `feedback_verify_tung_entry_truoc_khi_gui.md` ghi rõ '
  'HARD GATE: mọi task compile danh sách paper/IF tạp chí PHẢI verify '
  'từng entry bằng PubMed/Crossref/JCR TRƯỚC khi đưa vào file gửi đi. '
  'KHÔNG bao giờ dựa trí nhớ.')

P('Tổng cộng qua 3 vòng audit, em đã sửa: 9 lỗi IF/acceptance rate (v1→v2) '
  '+ 6 lỗi tác giả/năm/đối tượng (v2→v3) + 25 lỗi journal/author/IF (v3→v4) '
  '= 40 vấn đề trên 33 paper + 18 tạp chí.', italic=True)


# Clean metadata (sẽ được xử lý thêm bởi script xoa_metadata_va_watermark.py)
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(F)
print(f'Saved ERRATA-v3-to-v4 to: {F}')
print(f'Size: {os.path.getsize(F)} bytes')
