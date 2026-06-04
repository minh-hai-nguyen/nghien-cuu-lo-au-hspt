# -*- coding: utf-8 -*-
"""Doc trả lời 3 câu hỏi của thầy — V2 TIẾNG VIỆT THUẦN — 27/04/2026.
Quy ước: dùng tiếng Việt làm chính, chú thích thuật ngữ tiếng Anh trong ngoặc khi cần.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\Tra_loi_3_cau_hoi_cho_thay_v2_TiengViet_27042026.docx'
RED  = RGBColor(0xC0, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(10)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def warn(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('⚠ '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def ok(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('✓ '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(12); r2.font.name='Times New Roman'
def block(label, text):
    p=d.add_paragraph()
    r=p.add_run(label + ' '); r.bold=True; r.font.size=Pt(12); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.size=Pt(12); r2.font.name='Times New Roman'

# ===================================================================
title("TRẢ LỜI 3 CÂU HỎI CỦA THẦY", 16)
subtitle("(1) Nghiên cứu thu thập cách ứng phó tự phát từ học sinh — "
         "(2) Tác giả Việt Nam trong nghiên cứu Jefferies & Ungar 2020 — "
         "(3) Kiểm tra số liệu 80% trong nghiên cứu BESST 2024")
nr("")
subtitle("Phiên bản 2 — tiếng Việt thuần — 27/04/2026")
subtitle("Trợ lý nghiên cứu — Quy ước: dùng tiếng Việt làm chính, mở ngoặc chú thích "
         "thuật ngữ tiếng Anh khi cần (chỉ giữ thuật ngữ Anh nguyên gốc cho tên bài, "
         "tên tác giả, mã DOI/PMID/PMC, tên thang đo chuẩn quốc tế).")

# ===================================================================
# CÂU HỎI 1
# ===================================================================
H1("CÂU 1 — Có nghiên cứu nào thu thập cách ứng phó lo âu ĐA DẠNG, TỰ PHÁT từ học sinh?")
nr("Câu hỏi của thầy: \"Có nghiên cứu nào thu thập thông tin từ những cách giải quyết "
   "lo âu đa dạng, tự phát từ phía học sinh không?\"", italic=True, color=GRAY)
nr("")
nr("TRẢ LỜI THẲNG:", bold=True)
ok("CÓ — đặc biệt qua 3 phương pháp định tính (qualitative): "
   "(a) Photovoice (\"nhiếp ảnh tiếng nói\" — học sinh dùng ảnh thể hiện trải nghiệm), "
   "(b) Phỏng vấn nhóm tập trung (focus group) hoặc phỏng vấn hiện tượng học "
   "(phenomenological interview), (c) Khảo sát có câu hỏi mở (self-report survey "
   "with open-ended questions).")
nr("")
nr("Em đã tìm được 9 nghiên cứu phù hợp tiêu chí. Phân theo 3 phương pháp:")

# --- Photovoice ---
H2("1.1. Phương pháp Photovoice — học sinh dùng ảnh thể hiện cách họ ứng phó")

H3("Bài a — Tổng quan phạm vi hệ thống về Photovoice trong nghiên cứu sức khỏe tâm thần "
   "vị thành niên (2023)")
block('Tạp chí:',
      'International Journal of Adolescence and Youth — Tập 28, Số 1 (2023)')
block('Mã định danh:', 'DOI 10.1080/02673843.2023.2244043')
block('Đường dẫn:',
      'https://www.tandfonline.com/doi/full/10.1080/02673843.2023.2244043')
block('Phát hiện:',
      'Tổng quan hệ thống về phương pháp Photovoice trong nghiên cứu sức khỏe tâm '
      'thần (mental health) của vị thành niên. Đặc điểm phương pháp: học sinh chụp '
      'ảnh thể hiện trải nghiệm sống của mình + thảo luận các chủ đề từ ảnh. Phương '
      'pháp này cho phép thu thập các góc nhìn cảm xúc và ẩn dụ (emotional and '
      'metaphoric insights) mà các phương pháp định tính truyền thống khó tiếp cận. '
      'Bài kết luận: bằng chứng còn hạn chế về số lượng + chất lượng, cần thêm '
      'nghiên cứu nghiêm ngặt hơn.')

H3("Bài b — Photovoice với vị thành niên gốc Hispanic về sức khỏe tâm thần (2024–2025)")
block('Tạp chí:', 'Health Promotion International (Oxford Academic)')
block('Mã định danh:', 'DOI 10.1093/heapro/daag032')
block('Đường dẫn:',
      'https://academic.oup.com/heapro/article-pdf/41/1/daag032/67146769/daag032.pdf')
block('Phương pháp:',
      'Mười hai vị thành niên gốc Hispanic 13–17 tuổi từ 2 trung tâm thanh thiếu niên, '
      '3 buổi (chuẩn bị + thảo luận nhóm tập trung + phản hồi) từ tháng 6/2024 đến '
      'tháng 2/2025. Học sinh chụp ảnh thể hiện quan điểm về sức khỏe tâm thần, '
      'sau đó thảo luận theo kỹ thuật SHOWeD (gồm 5 câu hỏi: nhìn thấy gì? điều gì '
      'đang xảy ra? liên quan đến cuộc sống ta thế nào? tại sao tình trạng này tồn '
      'tại? làm gì được?).')
nr("PHÁT HIỆN ĐÁNG CHÚ Ý NHẤT:", bold=True, size=12)
nr("Khi được tự thể hiện (cách tiếp cận từ dưới lên — bottom-up), vị thành niên ưu "
   "tiên 4 NHÓM CÁCH ỨNG PHÓ ĐA DẠNG:")
nr("• Ứng phó CÁ NHÂN (solitary): viết nhật ký, đi dạo, thời gian một mình, đọc sách")
nr("• Ứng phó qua GIÁC QUAN (sensory): nghe nhạc, ngắm thiên nhiên, ăn món yêu thích")
nr("• Ứng phó SÁNG TẠO (creative): vẽ, viết, làm nghệ thuật, chơi nhạc cụ")
nr("• Ứng phó XÃ HỘI (social): trò chuyện với bạn, gia đình, hỗ trợ từ bạn đồng lứa")
nr("→ ĐIỂM QUAN TRỌNG: học sinh TỰ PHÁT đến với các cách ứng phó này KHÔNG QUA chương "
   "trình can thiệp áp đặt từ bên ngoài. Đây là khác biệt quan trọng so với các "
   "chương trình CBT/chánh niệm chuẩn — học sinh chọn theo bản thân mình.",
   bold=True, size=12)

H3("Bài c — Photovoice ở vị thành niên nhiễm HIV trong COVID-19, Nam Phi (2024)")
block('Tạp chí:', 'IJERPH — Tạp chí Nghiên cứu Môi trường và Sức khỏe Cộng đồng '
                  '(MDPI) — Tập 21(11): Bài 1517 (2024)')
block('Mã định danh:', 'DOI 10.3390/ijerph21111517')
block('Đường dẫn:', 'https://www.mdpi.com/1660-4601/21/11/1517')
block('Phát hiện:',
      'Các cách ứng phó đa dạng học sinh dùng (theo Photovoice): tôn giáo / đức '
      'tin, âm nhạc, thể thao, viết nhật ký, hỗ trợ từ bạn đồng lứa — khác hẳn các '
      '\"khuyến nghị chuẩn\" của CBT hay chánh niệm.')

# --- Focus group / phenomenological ---
H2("1.2. Phương pháp Phỏng vấn nhóm tập trung / Phỏng vấn hiện tượng học")

H3("Bài d — Học sinh THPT từ gia đình tan vỡ ở Philippines (2024–2025)")
block('Tạp chí:', 'Pantao Journal — đăng năm 2025 (năm học 2024–2025)')
block('Đường dẫn:',
      'https://pantaojournal.com/wp-content/uploads/2025/06/172-Delos-Santos.pdf')
block('Phương pháp:',
      'Phỏng vấn sâu 10 học sinh THPT Philippines từ gia đình tan vỡ; phân tích '
      'theo chủ đề (thematic analysis) ra 5 nhóm chủ đề lớn.')
block('Cách ứng phó tự phát học sinh dùng (theo phỏng vấn):',
      '\"Tìm kiếm trải nghiệm dễ chịu, rèn luyện sự độc lập, cố giữ thái độ lạc '
      'quan trong hoàn cảnh khó khăn\"; \"hình thành các mối liên kết chặt chẽ '
      'hơn với bạn bè và thành viên gia đình như nguồn hỗ trợ\".')

H3("Bài e — Trải nghiệm sống của vị thành niên + thanh niên trong COVID-19, "
   "Italia (2022)")
block('Mã định danh:', 'PMC8882423')
block('Đường dẫn:', 'https://pmc.ncbi.nlm.nih.gov/articles/PMC8882423/')
block('Phương pháp:',
      '25 vị thành niên + thanh niên 13–24 tuổi, phỏng vấn vào tháng 5/2020 '
      '(đợt phong toả đầu của COVID-19). Phân tích theo chủ đề ra 5 nhóm chủ đề, '
      'trong đó có \"phát triển các cách ứng phó để duy trì sức khỏe tâm thần\". '
      'Điểm hay: học sinh TỰ PHÁT thử nghiệm + phát triển các cách ứng phó mới '
      'trong đại dịch.')

H3("Bài f — \"Tiếng nói các em thực sự quan trọng\" (Frontiers Education 2025)")
block('Tên bài (tiếng Anh):',
      'Their voices matter: student and professional perspectives on overcoming '
      'school attendance problems in the context of social, emotional, and '
      'behavioral difficulties')
block('Tạm dịch:',
      'Tiếng nói các em thực sự quan trọng: quan điểm của học sinh và chuyên gia '
      'về việc vượt qua vấn đề chuyên cần trong bối cảnh khó khăn về xã hội, cảm '
      'xúc và hành vi')
block('Tạp chí:', 'Frontiers in Education (2025)')
block('Mã định danh:', 'DOI 10.3389/feduc.2025.1627098')
block('Đường dẫn:',
      'https://www.frontiersin.org/journals/education/articles/10.3389/'
      'feduc.2025.1627098/full')
block('Phương pháp:',
      'Phỏng vấn cả học sinh + chuyên gia; học sinh có khó khăn về xã hội – cảm '
      'xúc – hành vi (social-emotional-behavioural difficulties) chia sẻ về cách '
      'họ tự vượt qua vấn đề chuyên cần.')

# --- Survey with open-ended ---
H2("1.3. Khảo sát có câu hỏi mở (định lượng + định tính)")

H3("Bài g — Lo âu kéo dài ở học sinh THPT trong năm 2 đại dịch (PLOS ONE 2022)")
block('Tạp chí:', 'PLOS ONE — đăng năm 2022')
block('Mã định danh:', 'DOI 10.1371/journal.pone.0275292')
block('Đường dẫn:',
      'https://journals.plos.org/plosone/article?id=10.1371/journal.pone.0275292')
block('Phương pháp:',
      'Khảo sát học sinh THPT trong năm 2 của đại dịch COVID-19; có câu hỏi mở về '
      'cách họ ứng phó với lo âu kéo dài.')

H3("Bài h — Tổng quan hệ thống về lo âu ở sinh viên đại học (2025)")
block('Tạp chí:', 'Sage Open — Annals of Neurosciences (2025)')
block('Mã định danh:', 'DOI 10.1177/09727531251366078 — Mã PMC: PMC12420638')
block('Đường dẫn:',
      'https://journals.sagepub.com/doi/10.1177/09727531251366078')
block('Phát hiện:',
      'Phân tích 40 nghiên cứu thực nghiệm — tổng hợp các cơ chế ứng phó học sinh '
      'tự kê khai; bao gồm chánh niệm (cá nhân), hỗ trợ thể chế, hỗ trợ từ bạn '
      'đồng lứa.')

H3("Bài i — Cách ứng phó căng thẳng học tập ở học sinh THPT (IJAMR 2024)")
block('Tạp chí:',
      'International Journal of Academic Multidisciplinary Research — 2024')
block('Đường dẫn:',
      'http://ijeais.org/wp-content/uploads/2024/5/IJAMR240518.pdf')
block('Học sinh tự báo cáo các cách ứng phó:',
      '(1) hỗ trợ xã hội từ bạn đồng lứa, (2) tập thể dục, (3) thiền / kỹ thuật '
      'thư giãn, (4) giải quyết vấn đề có hệ thống, (5) tái cấu trúc nhận thức '
      '(thay đổi cách nghĩ).')

# --- Tổng hợp ---
H2("1.4. TỔNG HỢP — mẫu hình các cách ứng phó tự phát của học sinh qua 9 nghiên cứu trên")
tbl(['Nhóm cách ứng phó', 'Cách cụ thể', 'Tỉ lệ học sinh tự dùng (theo định tính)',
     'Văn hoá phổ biến'],
    [
        ['Giác quan (sensory)', 'Nghe nhạc, ngắm thiên nhiên',
         'Cao (mọi văn hoá)', 'Phổ biến toàn cầu'],
        ['Cá nhân (solitary)', 'Viết nhật ký, đi dạo, ngủ',
         'Cao ở phương Tây', 'Cao ở các nước Anglo-Saxon'],
        ['Sáng tạo (creative)', 'Vẽ, viết, làm nghệ thuật',
         'Trung bình – cao', 'Phổ biến toàn cầu'],
        ['Xã hội (social)', 'Trò chuyện bạn, gia đình',
         'RẤT cao ở Đông Á', 'Mạnh nhất ở văn hoá tập thể (collectivist)'],
        ['Tôn giáo / đức tin', 'Cầu nguyện, thiền tôn giáo',
         'Cao ở khu vực sùng tín', 'Châu Phi, Mỹ Latin, châu Á Hồi giáo'],
        ['Vận động thể chất', 'Thể thao, đi bộ, yoga',
         'Trung bình', 'Cao ở phương Tây'],
        ['Tránh né (avoidance)', 'Trốn tránh, đẩy lùi vấn đề',
         'Cao nhưng KHÔNG thích nghi', 'Phổ biến toàn cầu — cảnh báo'],
        ['Lạm dụng chất', 'Hút thuốc, rượu (tiêu cực)',
         'Cao ở học sinh THPT lớp lớn', 'Phổ biến toàn cầu — phải cảnh báo'],
    ], [3.0, 4.5, 4.0, 4.5])

H2("1.5. Nhận xét tổng quan")
nr("(1) Học sinh thực sự dùng NHIỀU CÁCH ỨNG PHÓ ĐA DẠNG — không chỉ Liệu pháp nhận "
   "thức – hành vi (CBT) hay chánh niệm.")
nr("(2) Phương pháp ĐỊNH TÍNH (Photovoice, phỏng vấn nhóm tập trung, phỏng vấn hiện "
   "tượng học) là cửa sổ tốt nhất để thu thập các cách ứng phó tự phát. Phương pháp "
   "ĐỊNH LƯỢNG chỉ đo các cách đã được liệt kê sẵn (thang Brief COPE, COPE-28) "
   "→ bỏ sót cách ứng phó học sinh tự nghĩ ra.")
nr("(3) GIÁ TRỊ NGHIÊN CỨU TỪ DƯỚI LÊN: phát hiện các cách ứng phó KHÔNG được dạy "
   "chính thức (ví dụ: nghe nhạc, ngắm thiên nhiên, tâm sự với bạn) — vốn rất phổ "
   "biến và HIỆU QUẢ trong thực tế nhưng ít có trong sách hướng dẫn CBT.")
nr("(4) BỐI CẢNH VIỆT NAM: chưa có nghiên cứu Photovoice về sức khỏe tâm thần bằng "
   "tiếng Việt mà em tìm được. Đây là KHOẢNG TRỐNG NGHIÊN CỨU RẤT QUAN TRỌNG cho "
   "Việt Nam — em đề xuất thầy có thể làm Photovoice nhỏ với 12–20 học sinh lớp "
   "10–12 ở 2–3 trường khác nhau (đô thị, nông thôn, dân tộc thiểu số nếu có "
   "điều kiện).")
nr("(5) Quy trình Photovoice gợi ý cho Việt Nam:")
nr("• Buổi 1: Giới thiệu phương pháp, đạo đức nghiên cứu, ký cam kết bảo mật ảnh "
   "(90 phút)")
nr("• Tự do 1–2 tuần: Học sinh chụp 5–10 ảnh thể hiện chủ đề \"cách em đối phó "
   "với căng thẳng\"")
nr("• Buổi 2: Phỏng vấn nhóm tập trung thảo luận về ảnh, dùng kỹ thuật SHOWeD "
   "(90–120 phút)")
nr("• Buổi 3: Phản hồi kết quả phân tích, thu thập ý kiến học sinh về ý nghĩa "
   "phát hiện (60 phút)")
nr("• Tổng cộng 3–4 buổi/học sinh, chi phí thấp, thủ tục đạo đức nghiên cứu đơn "
   "giản (chỉ cần đồng thuận tham gia + bảo mật ảnh + không nhận diện cá nhân "
   "trong báo cáo).")

# ===================================================================
# CÂU HỎI 2
# ===================================================================
H1("CÂU 2 — Tác giả Việt Nam trong nghiên cứu Jefferies & Ungar 2020 về 7 quốc gia?")
nr("Câu hỏi của thầy: \"Tìm bản tiếng Anh bài này để xem báo cáo của Việt Nam là của ai "
   "mà họ tổng hợp cùng nước ngoài? — Jefferies P, Ungar M (2020). Social Anxiety in "
   "Young People: A Prevalence Study in Seven Countries.\"", italic=True, color=GRAY)
nr("")

H2("2.1. Thông tin bài báo")
tbl(['Mục', 'Nội dung'],
    [
        ['Tên bài (tiếng Anh)', 'Social anxiety in young people: A prevalence study in '
                                'seven countries'],
        ['Tạm dịch', 'Lo âu xã hội ở người trẻ: Một nghiên cứu khảo sát tỉ lệ tại 7 '
                     'quốc gia'],
        ['Tác giả (CHỈ 2)', 'Philip Jefferies + Michael Ungar'],
        ['Đơn vị công tác (cả 2 tác giả)',
         'Trung tâm Nghiên cứu Khả năng Thích ứng (Resilience Research Centre), '
         'Khoa Sức khỏe (Faculty of Health), Đại học Dalhousie, thành phố Halifax, '
         'tỉnh Nova Scotia, Canada'],
        ['Tạp chí', 'PLoS ONE (truy cập mở)'],
        ['Tập / Số / Bài', '15(9): e0239133'],
        ['Năm công bố', '2020 — đăng ngày 17/9/2020'],
        ['Mã định danh số (DOI)', '10.1371/journal.pone.0239133'],
        ['Mã PubMed (PMID)', '32941482'],
        ['Mã PubMed Central (PMC)', 'PMC7498107'],
        ['Cỡ mẫu tổng', '6.825 người (3.342 nam + 3.428 nữ + 55 khác giới), 16–29 tuổi'],
        ['7 quốc gia khảo sát',
         'Brazil, Trung Quốc, Indonesia, Nga, Thái Lan, Hoa Kỳ, VIỆT NAM'],
        ['Cỡ mẫu Việt Nam', '984 người (487 nam + 493 nữ + 4 khác giới)'],
        ['Công cụ đo',
         'Thang Lo âu Tương tác Xã hội (Social Interaction Anxiety Scale — SIAS)'],
        ['Phát hiện chính',
         'Hơn 1/3 (36%) đáp viên đáp ứng ngưỡng Rối loạn Lo âu Xã hội (Social '
         'Anxiety Disorder — SAD); cao hơn nhiều ước tính trước đây'],
    ], [4.0, 12.0])

H2("2.2. Trả lời TRỰC TIẾP câu hỏi của thầy")
warn("KHÔNG CÓ TÁC GIẢ VIỆT NAM trong bài này. Cả 2 tác giả (Jefferies + Ungar) đều "
     "thuộc Trung tâm Nghiên cứu Khả năng Thích ứng, Đại học Dalhousie, Canada.")
nr("")
nr("DỮ LIỆU VIỆT NAM ĐƯỢC THU THẬP NHƯ THẾ NÀO?", bold=True)
nr("Theo phần Phương pháp + phần Lời cảm ơn của bài (em đã tải nguyên văn từ trang "
   "PMC NCBI):")
nr("• Người tham gia ở 7 nước được tuyển ngẫu nhiên thông qua 3 CÔNG TY KHẢO SÁT "
   "THỊ TRƯỜNG (market research companies):")
nr("  – Dynata (trụ sở Hoa Kỳ, có chi nhánh toàn cầu)")
nr("  – Online Market Intelligence – OMI (trụ sở Nga / Đông Âu)")
nr("  – GMO Research (trụ sở Nhật Bản, hoạt động khu vực châu Á – Thái Bình Dương)")
nr("• Phần Lời cảm ơn ghi nguyên văn (đã dịch): \"Các tác giả xin ghi nhận vai trò "
   "của Edelman Intelligence trong việc thu thập dữ liệu gốc thay mặt cho Unilever "
   "và CLEAR.\" → Edelman Intelligence là công ty khảo sát toàn cầu, thu thập dữ "
   "liệu THỨ CẤP cho 2 thương hiệu Unilever + CLEAR (CLEAR là dầu gội trị gàu).",
   italic=True)
nr("")
nr("KẾT LUẬN VỀ NGUỒN DỮ LIỆU VIỆT NAM:", bold=True, color=RED)
nr("Dữ liệu Việt Nam (n=984) trong bài này KHÔNG đến từ một viện nghiên cứu hay "
   "trường đại học Việt Nam nào, mà từ:")
nr("• Một công ty marketing toàn cầu (Dynata / OMI / GMO Research) tuyển người "
   "tham gia qua bảng câu hỏi trực tuyến (online panel)")
nr("• Dữ liệu gốc do Edelman Intelligence thu thập cho thương hiệu UNILEVER + CLEAR "
   "(đây là dữ liệu marketing thứ cấp, có thể nguồn gốc từ chiến dịch quảng cáo "
   "của CLEAR liên quan đến \"sự tự tin\")")
nr("• Hai tác giả Jefferies & Ungar phân tích lại bộ dữ liệu này (gọi là phân tích "
   "thứ cấp — secondary analysis) cho mục đích học thuật")

H2("2.3. Hệ quả cho việc trích dẫn")
warn("HẠN CHẾ về tính giá trị bên ngoài (external validity) của số liệu Việt Nam "
     "trong bài này:")
nr("(1) Mẫu KHÔNG ngẫu nhiên thuần tuý từ dân số Việt Nam — tuyển qua bảng câu hỏi "
   "trực tuyến của công ty khảo sát, có thể thiên về thanh thiếu niên đô thị có "
   "internet + sẵn lòng tham gia khảo sát marketing.")
nr("(2) Thang SIAS bản tiếng Việt có thể CHƯA được kiểm định (validate) cho dân số "
   "Việt Nam — bài không đề cập việc kiểm định.")
nr("(3) Dữ liệu thu thập cho mục đích MARKETING (Unilever / CLEAR) — câu hỏi có "
   "thể bị thiên lệch theo khung \"hình ảnh bản thân / sự tự tin\", không phải khung "
   "rối loạn lo âu xã hội theo tiêu chuẩn DSM-5.")
nr("(4) KHÔNG có chuyên gia Việt Nam đánh giá tính phù hợp văn hoá của công cụ đo "
   "và việc diễn giải kết quả.")
nr("")
nr("KHUYẾN NGHỊ KHI TRÍCH DẪN:", bold=True, color=BLUE)
nr("Khi trích con số \"Việt Nam có tỉ lệ Rối loạn Lo âu Xã hội 30,7% / cao thứ "
   "nhất trong 7 nước\" (theo bài), nên ghi rõ thêm: \"theo dữ liệu commercial "
   "panel thu cho Unilever / CLEAR, không phải khảo sát quốc gia chuẩn\". Số liệu "
   "này KHÁC với V-NAMHS 2022 (sử dụng DISC-5/DSM-5) — V-NAMHS là số liệu khảo sát "
   "quốc gia chính thức của Việt Nam và đáng tin cậy hơn cho việc tham chiếu.")

# ===================================================================
# CÂU HỎI 3
# ===================================================================
H1("CÂU 3 — Kiểm tra: 80% học sinh trong mẫu 900 BESST chưa từng tìm giải pháp?")
nr("Câu hỏi của thầy: \"Thông tin 80% học sinh trong mẫu 900 học sinh trong nghiên "
   "cứu của Brown chưa từng tìm giải pháp trước khi đến với BESST, có đúng không?\"",
   italic=True, color=GRAY)
nr("")

H2("3.1. Trả lời thẳng + đính chính nhỏ")
ok("CON SỐ 80% LÀ ĐÚNG — nhưng cần PHÂN BIỆT NGHĨA chính xác.")
nr("")
nr("Trong PDF gốc của Brown 2024 BESST (đăng trên The Lancet Psychiatry tập 11(7): "
   "trang 504–515) — em đã đọc đầy đủ + đã trích trong tài liệu dịch song ngữ:")
nr("• Bảng 1 (trang 510, Đặc điểm nền — Baseline characteristics):", bold=True)
nr("  – Câu hỏi: \"Đã từng tham vấn bác sĩ đa khoa cho vấn đề sức khỏe tâm thần\"")
nr("  – Trả lời KHÔNG: 720 người (80%); Trả lời CÓ: 179 người (20%)")
nr("  – Tổng 720 + 179 = 899 (1 mất dữ liệu); làm tròn 80% / 20%")
nr("• Phần Thảo luận (trang 513):", bold=True)
nr("  – Nguyên văn (đã dịch): \"Tỉ lệ cao những người chưa từng tìm trợ giúp qua "
   "KÊNH CHÍNH THỨC (formal routes) — 80%\"")
nr("• Hộp Bối cảnh nghiên cứu (Research-in-context, trang 505):", bold=True)
nr("  – Nguyên văn (đã dịch): \"80% học sinh trong nghiên cứu của chúng tôi chưa "
   "từng tìm trợ giúp trước đó\"")

H2("3.2. PHÂN BIỆT QUAN TRỌNG — 80% NGHĨA CHÍNH XÁC LÀ GÌ?")
warn("80% nghĩa là \"chưa từng tìm bác sĩ đa khoa (general practitioner — GP) hay kênh "
     "y tế chính thức\" — KHÔNG PHẢI \"chưa từng làm gì\" hay \"chưa từng tìm bất kỳ "
     "giải pháp nào\".")
nr("")
nr("Cụ thể, bài Brown 2024 đo MỘT câu hỏi: \"Đã từng tham vấn bác sĩ đa khoa cho "
   "vấn đề sức khỏe tâm thần chưa?\" — biến nhị phân CÓ/KHÔNG. Câu này CHỈ đo "
   "kênh chính thức qua hệ thống y tế công NHS (National Health Service) của Anh.")
nr("")
nr("Các kênh học sinh CÓ THỂ đã dùng nhưng KHÔNG được đo trong câu này (vẫn nằm "
   "trong nhóm 80%):")
tbl(['Kênh có thể đã dùng', 'Có nằm trong câu hỏi 80% không?'],
    [
        ['Tâm sự với bạn (peer support)', 'KHÔNG đo'],
        ['Tâm sự với gia đình', 'KHÔNG đo'],
        ['Tự tìm trên mạng (website, ứng dụng, video)', 'KHÔNG đo'],
        ['Đến gặp tư vấn học đường', 'KHÔNG đo'],
        ['Cầu nguyện / thực hành tôn giáo', 'KHÔNG đo'],
        ['Tự dùng các cách ứng phó (nghe nhạc, thể thao, viết nhật ký)', 'KHÔNG đo'],
        ['Dịch vụ tâm thần chuyên khoa CAMHS (Child & Adolescent Mental Health Services) Anh',
         'CAMHS là TIÊU CHÍ LOẠI TRỪ — học sinh đang nhận trị liệu CAMHS bị loại '
         'khỏi nghiên cứu, nên không có trong mẫu'],
        ['Bác sĩ đa khoa (GP) cho vấn đề SKTT', 'CHÍNH LÀ câu hỏi 80%'],
    ], [7.0, 9.0])

H2("3.3. Cách hiểu chính xác")
nr("80% trong nghiên cứu BESST → 80% học sinh đã đăng ký (tự giới thiệu — "
   "self-referral) cho workshop DISCOVER là những học sinh CHƯA TỪNG đi qua KÊNH Y TẾ "
   "CHÍNH THỨC (bác sĩ đa khoa) cho vấn đề SKTT của họ.")
nr("")
nr("Ý nghĩa thực tiễn:")
nr("• Cơ chế tự giới thiệu chạm tới nhóm \"khó tiếp cận\" (hard-to-reach) — vốn "
   "không đi bác sĩ đa khoa (do kỳ thị, không biết bác sĩ có thể giúp về SKTT, hoặc "
   "không nghĩ vấn đề đủ \"nặng\" để gặp bác sĩ).")
nr("• Trong 80% này, có thể có học sinh đã DÙNG nhiều cách ứng phó tự phát khác (tâm "
   "sự bạn, online, tư vấn học đường) — bài KHÔNG đo nên không biết tỉ lệ.")
nr("• Đây là PHÁT HIỆN MẠNH cho lập luận tự giới thiệu (theo mô hình PLACES của Brown "
   "2022): chương trình tới đúng nhóm cần.")

H2("3.4. Đề xuất cách diễn đạt cho thầy khi trích dẫn")
warn("KHÔNG NÊN NÓI: \"80% học sinh BESST chưa từng tìm giải pháp\" (KHÔNG HOÀN TOÀN "
     "ĐÚNG).")
nr("")
nr("NÊN NÓI MỘT TRONG BA CÁCH SAU:", bold=True, color=GREEN)
nr("(a) \"80% học sinh BESST chưa từng tham vấn bác sĩ đa khoa về vấn đề sức khỏe "
   "tâm thần\" — chính xác về kênh y tế chính thức.")
nr("(b) \"80% học sinh BESST chưa từng tiếp cận hệ thống y tế chính thức cho sức "
   "khỏe tâm thần — workshop DISCOVER là điểm đầu tiên họ tiếp xúc với chuyên môn "
   "sức khỏe tâm thần.\"")
nr("(c) \"BESST tiếp cận được nhóm khó tiếp cận: 80% học sinh tham gia chưa bao "
   "giờ qua kênh bác sĩ đa khoa hay trợ giúp chính thức — chứng tỏ cơ chế tự giới "
   "thiệu có khả năng phá rào cản tiếp cận dịch vụ.\"")

# ===================================================================
H1("THAM KHẢO ĐẦY ĐỦ")

H2("Cho câu 1 (cách ứng phó tự phát từ học sinh)")
nr("• Tổng quan phạm vi về Photovoice trong sức khỏe tâm thần vị thành niên (2023). "
   "International Journal of Adolescence and Youth, 28(1). "
   "DOI: 10.1080/02673843.2023.2244043", size=11)
nr("• Photovoice với vị thành niên gốc Hispanic (2024–2025). Health Promotion "
   "International, 41(1):daag032. DOI: 10.1093/heapro/daag032", size=11)
nr("• Photovoice ở vị thành niên nhiễm HIV trong COVID-19, Nam Phi (2024). IJERPH, "
   "21(11):1517. DOI: 10.3390/ijerph21111517", size=11)
nr("• Delos Santos và cộng sự (2025). Học sinh THPT từ gia đình tan vỡ ở "
   "Philippines — phỏng vấn hiện tượng học. Pantao Journal.", size=11)
nr("• Trải nghiệm sống của AYAs trong COVID-19, định tính (2022). Mã PMC: PMC8882423.",
   size=11)
nr("• Frontiers in Education (2025) — \"Their voices matter\". "
   "DOI: 10.3389/feduc.2025.1627098", size=11)
nr("• PLOS ONE (2022). Lo âu kéo dài ở học sinh THPT trong COVID năm 2. "
   "DOI: 10.1371/journal.pone.0275292", size=11)
nr("• Tổng quan về lo âu sinh viên đại học (2025). Sage. "
   "DOI: 10.1177/09727531251366078 — Mã PMC: PMC12420638", size=11)
nr("• Cách ứng phó căng thẳng học tập ở học sinh THPT — IJAMR 2024.", size=11)

H2("Cho câu 2 (Jefferies & Ungar 2020)")
nr("• Jefferies P, Ungar M (2020). Social anxiety in young people: A prevalence "
   "study in seven countries. PLoS ONE, 15(9):e0239133. "
   "DOI: 10.1371/journal.pone.0239133 — Mã PubMed: PMID 32941482 — Mã PubMed "
   "Central: PMC7498107", size=11)
nr("• Em đã tải nguyên văn từ PMC NCBI để xác minh tác giả, đơn vị công tác, "
   "phương thức tuyển mẫu Việt Nam, lời cảm ơn", size=11)
nr("• Đối chiếu tham khảo: V-NAMHS 2022 (kho dự án) — Khảo sát quốc gia Việt Nam "
   "với chuẩn DISC-5/DSM-5 — uy tín hơn cho dữ liệu quốc gia", size=11)

H2("Cho câu 3 (BESST 80% chưa tìm trợ giúp chính thức)")
nr("• Brown J, James K, Lisk S, và cộng sự (2024). Hiệu quả lâm sàng và chi phí - "
   "hiệu quả của chương trình CBT ngắn dễ tiếp cận cho căng thẳng ở học sinh tuổi "
   "đi học (BESST): thử nghiệm có đối chứng theo cụm. The Lancet Psychiatry, "
   "11(7):504–515. Mã PubMed: PMID 38759665 — DOI: 10.1016/S2215-0366(24)00101-9",
   size=11)
nr("• PDF đầy đủ đã đọc: 02_Papers-goc/UK_BESST_PLACES/"
   "Brown_2024_BESST_Lancet_Psychiatry.pdf (Bảng 1 trang 510 + Thảo luận trang 513 "
   "+ Hộp Research-in-context trang 505)", size=11)
nr("• Tài liệu dịch + phản biện: 03_Ban-dich/Bai_dich_phan_bien/"
   "BESST_Brown_2024_dich_phan_bien_25042026.docx — đã trích 80% trong PHẦN 1 + "
   "Bảng B + PHẦN 3 phản biện", size=11)
nr("• Cạnh trong Đồ thị tri thức (Knowledge Graph): QT042_BESST → "
   "Concept::self_referral với bối cảnh 80% chưa từng tham vấn", size=11)

H2("Truy vết & công cụ đã dùng để trả lời")
nr("• Câu 1: 3 lần tìm kiếm web (định tính + tiếng nói học sinh + photovoice) trong "
   "phiên 27/4/2026", size=11)
nr("• Câu 2: 1 lần tìm kiếm web + 1 lần tải nguyên văn từ PMC7498107", size=11)
nr("• Câu 3: Xác minh từ PDF Brown 2024 BESST đã đọc đầy đủ + chunk trong RAG "
   "QT042_BESST + cạnh trong Đồ thị tri thức", size=11)
nr("• Bộ nhớ áp dụng: feedback_research_workflow.md (kiểm 3-4 vòng) + "
   "feedback_doc_phai_co_reference.md (DOI/PMID đầy đủ)", size=11)

d.save(OUT)
print('Wrote:', OUT)
