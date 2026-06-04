# -*- coding: utf-8 -*-
"""
Phase A — Trích author từ 68 bài canonical (Tom-tat-tung-bai/ + PDF trang 1).

Output: 06_Scripts/author_kg_data/authors_raw.json
  {
    "QT001": {
      "paper_title": "...",
      "year": "2023",
      "journal": "...",
      "country": "USA",
      "authors_raw": "Jenkins K, Smith J, Doe A",
      "authors_list": [{"name": "Jenkins K", "position": 1, "is_corresponding": True}, ...]
    }
  }
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
PAPERS_ROOT = os.path.join(ROOT, '02_Papers-goc')
OUT_DIR = os.path.join(os.path.dirname(__file__), 'author_kg_data')
os.makedirs(OUT_DIR, exist_ok=True)

with open(os.path.join(PAPERS_ROOT, 'canonical_index.json'), encoding='utf-8') as f:
    canon = json.load(f)

def read_docx_text(path):
    if not os.path.exists(path): return ''
    try:
        d = Document(path)
        lines = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for r in t.rows:
                for c in r.cells:
                    lines.append(c.text.strip())
        return '\n'.join(lines)
    except:
        return ''

def extract_authors_from_summary(text, cid):
    """Try to extract author string from Tác giả field in summary."""
    # Pattern: "Tác giả: Author1, Author2..."
    # Common patterns in our Tom-tat-tung-bai files
    patterns = [
        r'T\u00e1c gi\u1ea3[:\s]+([^\n]+?)(?=\n[A-Z]|$)',  # Tác giả: ...
        r'authors?[:\s]+([^\n]+?)(?=\n|$)',  # English
        r'\b(?:của|do)\s+((?:[A-ZĐ][a-zà-ỹ]+(?:\s+[A-ZĐ][a-zà-ỹ]+){1,4},?\s*)+(?:và(?:\s+cộng\s+sự)?|et al\.?)?)\s*\(?\d{4}',
    ]
    for pat in patterns:
        m = re.search(pat, text, re.IGNORECASE | re.MULTILINE)
        if m:
            return m.group(1).strip()
    # Fallback: try first sentence with "2024"/"2023"/"2025" year near author names
    m = re.search(r'C\u00f4ng tr\u00ecnh\s+[\u00ab"\u201c]([^\u201d\u00bb"]+)[\u201d\u00bb"]\s+(?:của|do)\s+([^.]+?(?=\.))', text)
    if m:
        return m.group(2).strip()
    return ''

def parse_author_list(raw):
    """Parse raw author string into list. Handle 'et al', 'và cộng sự', separators."""
    if not raw: return []
    raw = raw.strip().rstrip('.').rstrip(',')
    # Strip trailing et al / cộng sự
    raw = re.sub(r',?\s*(?:et al\.?|và cộng sự|và cs\.?|và ĐTTT|và ĐTTV)\b.*', ' [et al.]', raw, flags=re.IGNORECASE)
    has_etal = '[et al.]' in raw
    raw = raw.replace(' [et al.]', '').strip()
    # Split by comma or '&' or 'và ' or ';'
    parts = re.split(r'\s*(?:,|&|;|\s+và\s+|\s+and\s+)\s*', raw)
    parts = [p.strip() for p in parts if p.strip() and len(p.strip()) > 1]
    authors = []
    for i, p in enumerate(parts):
        # Skip if looks like year / journal / too short
        if re.match(r'^\d+', p) or len(p) < 3:
            continue
        # Skip if no letter
        if not re.search(r'[a-zA-ZÀ-ỹ]', p):
            continue
        authors.append({
            'name': p,
            'position': i + 1,
            'is_corresponding': (i == 0),  # heuristic: first author
        })
    return authors, has_etal

# Main extraction
results = {}
stats = {'total': 0, 'with_authors': 0, 'empty': 0, 'papers_with_etal': 0}

for cid, meta in sorted(canon.items()):
    stats['total'] += 1
    summary_fn = meta.get('summary', '')
    pdf_fn = meta.get('pdf', '')
    descriptor = meta.get('descriptor', '')
    folder = meta.get('pdf_folder', '')

    summary_path = os.path.join(TT_DIR, summary_fn) if summary_fn else ''
    summary_text = read_docx_text(summary_path) if summary_path else ''

    authors_raw = extract_authors_from_summary(summary_text, cid)
    authors_list, has_etal = parse_author_list(authors_raw) if authors_raw else ([], False)

    if authors_list:
        stats['with_authors'] += 1
    else:
        stats['empty'] += 1
    if has_etal:
        stats['papers_with_etal'] += 1

    # Extract year from descriptor
    year_m = re.search(r'(\d{4})', descriptor)
    year = year_m.group(1) if year_m else ''

    # Country hint from descriptor / folder
    country = ''
    if folder == 'Viet-Nam' or 'VN' in cid:
        country = 'Vietnam'
    elif 'USA' in descriptor or 'America' in descriptor:
        country = 'USA'
    elif 'China' in descriptor or 'Chinese' in descriptor:
        country = 'China'
    elif 'Korea' in descriptor:
        country = 'Korea'
    elif 'Japan' in descriptor:
        country = 'Japan'
    elif 'Indonesia' in descriptor:
        country = 'Indonesia'
    elif 'India' in descriptor or 'Assam' in descriptor:
        country = 'India'
    elif 'SriLanka' in descriptor:
        country = 'Sri Lanka'
    elif 'Ireland' in descriptor:
        country = 'Ireland'
    elif 'Australia' in descriptor:
        country = 'Australia'
    elif 'UK' in descriptor or 'NHS' in descriptor:
        country = 'UK'
    elif 'Norway' in descriptor:
        country = 'Norway'
    elif 'Ethiopia' in descriptor:
        country = 'Ethiopia'
    elif 'Saudi' in descriptor:
        country = 'Saudi Arabia'
    elif 'Filipino' in descriptor or 'Philippines' in descriptor:
        country = 'Philippines'

    results[cid] = {
        'descriptor': descriptor,
        'year': year,
        'country_hint': country,
        'pdf_folder': folder,
        'authors_raw': authors_raw,
        'authors_list': authors_list,
        'has_etal': has_etal,
        'n_authors_extracted': len(authors_list),
        'summary_file': summary_fn,
        'pdf_file': pdf_fn,
    }

out_path = os.path.join(OUT_DIR, 'authors_raw.json')
with open(out_path, 'w', encoding='utf-8') as f:
    json.dump(results, f, ensure_ascii=False, indent=2)

print('='*70)
print('PHASE A — AUTHOR EXTRACTION RAW')
print('='*70)
print(f'Total papers processed: {stats["total"]}')
print(f'Papers with authors extracted: {stats["with_authors"]}')
print(f'Papers empty (need PDF fallback): {stats["empty"]}')
print(f'Papers with "et al" marker: {stats["papers_with_etal"]}')
print()
print(f'Output: {out_path}')

# Preview first 10
print()
print('Preview first 10:')
for cid in list(results.keys())[:10]:
    r = results[cid]
    n = r['n_authors_extracted']
    raw = r['authors_raw'][:100] if r['authors_raw'] else '(empty)'
    print(f'  {cid}: {n} authors — raw: {raw}')
