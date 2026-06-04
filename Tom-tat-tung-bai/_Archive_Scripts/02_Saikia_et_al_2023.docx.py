# -*- coding: utf-8 -*-
"""Bài 2: Saikia et al. 2023 — theo phong cách CTH v5"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style.paragraph_format.space_after = Pt(4)
style.paragraph_format.line_spacing = 1.5
for s in doc.sections:
    s.top_margin = Cm(2.5)
    s.bottom_margin = Cm(2.5)
    s.left_margin = Cm(3)
    s.right_margin = Cm(2)

def rb(text):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = RED

def bl(text, bold=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.font.color.rgb = BLUE; r.bold = bold

def rh2(text):
    h = doc.add_heading(text, level=2)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = RED

def bh3(text):
    h = doc.add_heading(text, level=3)
    for r in h.runs: r.font.name = 'Times New Roman'; r.font.color.rgb = BLUE

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr()
    we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa')
    tcW.append(we)

def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)):
            set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)

# ===== BÀI 2 =====
bl('Tóm tắt bài 2', bold=True)

# BƯỚC 1: ĐỊNH DANH (1 câu dài)
rb('Tên công trình, tác giả, năm, mẫu khảo sát, đối tượng, khách thể, địa bàn khảo sát')
bl('Công trình nghiên cứu \u00ab Các bệnh lý sức khỏe tâm thần và các yếu tố liên quan ở thanh thiếu niên tại Kamrup (Metro), Assam: Một nghiên cứu tại trường học \u00bb (Mental Health Morbidities and their Correlates among the Adolescents in Kamrup (Metro), Assam: A School-Based Study), do Anku M. Saikia, Hemen Das và Vinoth Rajendran (2023) tại Khoa Y học Cộng đồng, Trường Y khoa và Bệnh viện Gauhati, Guwahati, Assam, Ấn Độ, khảo sát trên mẫu 360 thanh thiếu niên từ 14 đến 17 tuổi tại 10 trường trung học công lập ở quận Kamrup (Metro), vùng Đông Bắc Ấn Độ.')

# BƯỚC 2: TỔNG QUAN PP + "Nói cách khác" (KT1)
rb('Phương pháp nghiên cứu')
bl('Công trình này sử dụng Thang đo Trầm cảm Lo âu Căng thẳng 21 mục (DASS-21 \u2014 Depression Anxiety Stress Scale) phiên bản tiếng Assam kết hợp với lịch phỏng vấn được thiết kế sẵn và thang đo Kuppuswamy sửa đổi để đánh giá tình trạng kinh tế xã hội. Nói cách khác, công trình này sử dụng phương pháp sàng lọc cắt ngang tại trường học.', bold=True)

# BƯỚC 3: ĐỊNH NGHĨA (KT5 — phả hệ, mỗi công cụ 1 đoạn)
bl('DASS-21 là phiên bản rút gọn 21 mục của Thang đo Trầm cảm Lo âu Căng thẳng (Depression Anxiety Stress Scale) gốc 42 mục (Lovibond & Lovibond, 1995), được thiết kế để sàng lọc đồng thời ba tình trạng: trầm cảm, lo âu và căng thẳng.')
bl('Thang đo Kuppuswamy sửa đổi (Modified Kuppuswamy Scale) là công cụ đánh giá tình trạng kinh tế xã hội được sử dụng rộng rãi tại Ấn Độ, dựa trên ba tiêu chí: học vấn của trụ cột gia đình, nghề nghiệp và thu nhập hộ gia đình.')

# BƯỚC 4: BIỆN MINH (KT4)
bl('Tổng quan tài liệu của chúng tôi cho thấy phần lớn các nghiên cứu về sức khỏe tâm thần thanh thiếu niên tại Ấn Độ tập trung vào các bang miền Nam và miền Bắc, trong khi vùng Đông Bắc Ấn Độ \u2014 nơi có đặc điểm văn hóa bộ lạc khác biệt \u2014 hầu như chưa được khảo sát. Đây là nghiên cứu tiên phong tại khu vực này.', bold=True)

# BƯỚC 5: LIỆT KÊ DỮ LIỆU
bl('Dữ liệu nhân khẩu xã hội bao gồm tuổi, giới tính, loại gia đình (hạt nhân hoặc mở rộng), tầng lớp kinh tế xã hội, tình trạng cha mẹ (đơn thân hoặc đủ), sử dụng rượu của cha mẹ, tiền sử bị lạm dụng, tiền sử bệnh tâm thần trong gia đình, thói quen chơi trò chơi điện tử và kết quả học tập.')

# BƯỚC 6: MÔ TẢ QUY TRÌNH
bl('Nghiên cứu được thực hiện từ tháng 4/2019 đến tháng 6/2020. Từ tổng số 120 trường trung học công lập, 10 trường được chọn ngẫu nhiên đơn giản. Từ mỗi trường, 36 học sinh được chọn ngẫu nhiên (12 từ mỗi lớp 8, 9, 10), đảm bảo số lượng nam và nữ bằng nhau. Quyền riêng tư và tính bảo mật nghiêm ngặt được duy trì trong quá trình phỏng vấn. DASS-21 được dịch sang tiếng Assam, dịch ngược bởi chuyên gia song ngữ và xác nhận bởi nhóm chuyên gia tâm thần học và y học cộng đồng.')

# BƯỚC 9: ĐỊNH LƯỢNG — BẢNG
rb('Kết quả nghiên cứu định lượng')

bl('Tỷ lệ trầm cảm, lo âu và căng thẳng theo DASS-21 (N = 360):', bold=True)
tbl(
    ['Tình trạng', 'Tỷ lệ', 'n'],
    [
        ['Trầm cảm', '22,2%', '80/360'],
        ['Lo âu', '24,4%', '88/360'],
        ['Căng thẳng', '6,9%', '25/360'],
    ],
    widths=[5.0, 3.0, 3.0]
)
doc.add_paragraph()

bl('Tỷ lệ lo âu theo giới tính:', bold=True)
tbl(
    ['Giới tính', 'Có lo âu', 'Không lo âu', 'P'],
    [
        ['Nam', '54 (30,0%)', '126 (70,0%)', ''],
        ['Nữ', '34 (18,9%)', '146 (81,1%)', '0,049'],
    ],
    widths=[3.0, 3.5, 3.5, 2.0]
)
doc.add_paragraph()

bl('Yếu tố liên quan có ý nghĩa thống kê với lo âu:', bold=True)
tbl(
    ['Yếu tố', 'P (lo âu)', 'P (trầm cảm)', 'Ghi chú'],
    [
        ['Tình trạng cha mẹ (đơn thân)', '<0,001', '0,001', '52,6\u201363,2% nếu đơn thân'],
        ['Cha mẹ sử dụng rượu', '<0,001', '<0,001', '38,3\u201340,4% nếu có rượu'],
        ['Chơi trò chơi điện tử', '0,003', '<0,001', '27,9\u201330,8% nếu chơi game'],
        ['Kết quả học tập (lưu ban)', '<0,001', '<0,001', '51,7\u201358,6% nếu lưu ban'],
        ['Tầng lớp KT-XH thấp', '0,028', '0,839', 'Chỉ có ý nghĩa với lo âu'],
        ['Tiền sử bị lạm dụng', '0,044', '0,076', 'Chỉ có ý nghĩa với lo âu'],
        ['Tiền sử bệnh TT gia đình', '0,043', '0,471', 'Chỉ có ý nghĩa với lo âu'],
        ['Giới tính (nam > nữ)', '0,049', '0,076', 'Trái y văn quốc tế'],
    ],
    widths=[4.0, 2.0, 2.5, 4.0]
)
doc.add_paragraph()

# BƯỚC 10: ĐỐI CHIẾU (KT3 — xác nhận 3 bước)
bh3('Khác biệt giới tính : Nam giới luôn được báo cáo có tỷ lệ lo âu thấp hơn nữ giới trong y văn quốc tế (McLean và cộng sự, 2011). Thực tế, sự chênh lệch giới trong lo âu ở thanh thiếu niên đã được xác nhận trong hàng chục nghiên cứu đa quốc gia (Daly, 2022). Tuy nhiên, nghiên cứu hiện tại cho thấy kết quả TRÁI NGƯỢC: nam giới có tỷ lệ lo âu cao hơn nữ (30,0% so với 18,9%, P = 0,049). Phát hiện này phù hợp với Wen và cộng sự (2020) tại nông thôn Trung Quốc và Xu và cộng sự (2021) trong bối cảnh COVID-19 tại Trung Quốc.')

# BƯỚC 11+12: NHẬN XÉT + CẦU NỐI
rb('Nhận xét, phát hiện qua kết quả nghiên cứu')
bh3('Phát hiện về yếu tố gia đình, lối sống và bối cảnh văn hóa Đông Bắc Ấn Độ')
bl('*Cha mẹ đơn thân và sử dụng rượu.* Nổi bật đối với cả nam và nữ thanh thiếu niên là mối liên hệ mạnh giữa hoàn cảnh gia đình bất lợi và sức khỏe tâm thần. Thanh thiếu niên sống với cha hoặc mẹ đơn thân có tỷ lệ lo âu 63,2% \u2014 gấp gần 3 lần so với nhóm sống với cả hai bố mẹ (22,3%). Đây là một yếu tố nguy cơ không nằm trong các mục sàng lọc của DASS-21 nhưng được xác nhận rõ ràng bởi phân tích chi-square.')
bl('*Trò chơi điện tử và kết quả học tập.* Sự tác động của thói quen chơi trò chơi điện tử đến sức khỏe tâm thần thanh thiếu niên là rất nổi bật. Những vấn đề này được liên hệ với kết quả học tập kém \u2014 thanh thiếu niên bị lưu ban có tỷ lệ lo âu 51,7%, gấp đôi nhóm có kết quả bình thường.')
bl('*Bối cảnh văn hóa bộ lạc Đông Bắc.* Các phân tích cho thấy bối cảnh liên quan đến trải nghiệm lo âu của thanh thiếu niên, bao gồm đặc thù văn hóa bộ lạc vùng Đông Bắc Ấn Độ \u2014 nơi kỳ vọng về vai trò nam giới có thể khác biệt so với phần còn lại của Ấn Độ, gợi ý rằng yếu tố văn hóa địa phương có thể đảo ngược xu hướng giới tính thông thường.')

# BƯỚC 13: KẾT LUẬN (Kiến trúc B — bằng chứng chèn giữa)
rb('Kết luận')
bl('Dữ liệu của chúng tôi, cho thấy sự hiện diện của các tỷ lệ đáng kể trầm cảm (22,2%), lo âu (24,4%) và căng thẳng (6,9%) ở thanh thiếu niên tại trường học vùng Đông Bắc Ấn Độ \u2014 một khu vực hầu như chưa được nghiên cứu trước đây, gợi ý rằng cần chú ý đến các yếu tố gia đình (cha mẹ đơn thân, sử dụng rượu) và lối sống (chơi game, giấc ngủ) trong thiết kế can thiệp. Sự khác biệt giới tính ngược (nam > nữ), trước đây hiếm được báo cáo, đã có thể quan sát được trong nghiên cứu này và có ý nghĩa tiềm năng cho việc nhận diện các yếu tố văn hóa đặc thù khu vực. Trong bối cảnh trường học Đông Bắc Ấn Độ như nghiên cứu hiện tại, nhu cầu giải quyết mối quan hệ giữa các yếu tố gia đình, lối sống và sức khỏe tâm thần, dù ở mức nhẹ hay nặng, là thiết yếu để tạo điều kiện cho sự tham gia giáo dục.', bold=True)

# PHẢN BIỆN
rh2('Phản biện')
bl('Nam giới có tỷ lệ lo âu cao hơn nữ (30,0% so với 18,9%, P = 0,049) \u2014 trái ngược hoàn toàn với y văn quốc tế, nhưng tác giả không giải thích sâu cơ chế văn hóa cụ thể của vùng Đông Bắc. DASS-21 phiên bản tiếng Assam thiếu báo cáo Cronbach alpha \u2014 không rõ độ tin cậy của phiên bản dịch. Chỉ sử dụng chi-square (đơn biến) mà không dùng hồi quy logistic đa biến \u2014 không thể đánh giá ảnh hưởng độc lập của từng yếu tố sau khi kiểm soát biến gây nhiễu (Rothman và cộng sự, 2008). Thu thập dữ liệu từ tháng 4/2019 đến tháng 6/2020 \u2014 trùng với giai đoạn đầu đại dịch COVID-19 nhưng tác giả không đề cập đến tác động tiềm tàng của đại dịch lên kết quả (Racine và cộng sự, 2021). Tuy nhiên, đây là nghiên cứu tiên phong tại Đông Bắc Ấn Độ với thiết kế lấy mẫu ngẫu nhiên tốt từ 10/120 trường.')

# HƯỚNG NC
rh2('Hướng nghiên cứu tiếp theo')
bl('Nghiên cứu quy mô lớn hơn tại nhiều quận vùng Đông Bắc Ấn Độ với hồi quy logistic đa biến. Khảo sát sâu yếu tố văn hóa bộ lạc đặc thù giải thích tại sao nam giới lo âu cao hơn nữ. So sánh thành thị-nông thôn trong cùng khu vực. Nghiên cứu can thiệp tại trường học nhắm vào giảm thói quen chơi game và hỗ trợ gia đình đơn thân. Xác thực DASS-21 tiếng Assam với báo cáo Cronbach alpha đầy đủ (Lovibond & Lovibond, 1995).')

doc.add_paragraph()
p = doc.add_paragraph()
r = p.add_run('Đánh giá chất lượng: \u2b50\u2b50\u2b50 Trung bình. Nghiên cứu tiên phong cho Đông Bắc Ấn Độ với lấy mẫu ngẫu nhiên tốt, nhưng thiếu phân tích đa biến và thiếu giải thích cho phát hiện nam > nữ.')
r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(12)

doc.save('02_Saikia_et_al_2023.docx')
