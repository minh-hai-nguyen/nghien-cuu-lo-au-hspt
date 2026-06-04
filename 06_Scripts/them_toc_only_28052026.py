# -*- coding: utf-8 -*-
"""CHI THEM TOC - khong dong gi khac.
File: 01_LuanAn_v3_1_FixCoping_28052026.docx
28/05/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.text.paragraph import Paragraph

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

# Backup
backup = FILE.replace('.docx', '_BEFORE_TOC.docx')
shutil.copy(FILE, backup)
print(f"Backup: {backup}")

d = Document(FILE)
print(f"File: {FILE}")
print(f"Paragraphs: {len(d.paragraphs)}")

# Find MỤC LỤC
toc_idx = None
for i, p in enumerate(d.paragraphs[:200]):
    if p.text.strip().upper() in ('MỤC LỤC', 'MUC LUC'):
        toc_idx = i
        break

if toc_idx is None:
    print("KHONG TIM THAY MỤC LỤC!")
    sys.exit(1)

print(f"Found MỤC LỤC at para {toc_idx}")

# Check if TOC field already exists immediately after
already_has_toc = False
for elem in d.element.body.iter():
    if elem.tag == qn('w:instrText'):
        if elem.text and 'TOC' in elem.text:
            already_has_toc = True
            break

if already_has_toc:
    print("CANH BAO: TOC field da co san trong tai lieu. Khong chen them.")
    sys.exit(0)

# Insert TOC field
def make_toc_paragraph(parent_para):
    r = parent_para.add_run()
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fldChar_begin)
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r._r.append(instrText)
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fldChar_sep)
    instrResult = OxmlElement('w:t')
    instrResult.text = 'Nhấn F9 hoặc chuột phải > Update Field để cập nhật Mục lục.'
    r._r.append(instrResult)
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar_end)


toc_heading_p = d.paragraphs[toc_idx]
new_p_xml = OxmlElement('w:p')
toc_heading_p._element.addnext(new_p_xml)
toc_p_obj = Paragraph(new_p_xml, toc_heading_p._parent)
make_toc_paragraph(toc_p_obj)
print(f"Inserted TOC field after para {toc_idx} (MỤC LỤC)")

# Save back to same file
d.save(FILE)
print(f"\nSAVED: {FILE}")
print(f"Size: {os.path.getsize(FILE)//1024} KB")
print("\nNOTHING ELSE was changed.")
