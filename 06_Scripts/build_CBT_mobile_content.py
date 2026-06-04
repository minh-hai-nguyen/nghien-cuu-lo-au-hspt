"""Doc trả lời: Nội dung cụ thể chương trình CBT mobile cho VTN — Nhật, TQ, Anh."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

BLUE = RGBColor(0, 112, 192)
DARK = RGBColor(31, 73, 125)

def shade_cell(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)

def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color
    return h

def p_blue(text, bold=False):
    p = d.add_paragraph(); r = p.add_run(text); r.font.color.rgb = BLUE; r.font.size = Pt(11); r.bold = bold
    return p

# Title
title = d.add_heading('Nội dung chương trình CBT trên điện thoại cho thanh thiếu niên: Nhật, Trung Quốc, Anh', level=0)
for r in title.runs: r.font.color.rgb = DARK
sub = d.add_paragraph()
sr = sub.add_run('Tổng hợp 4 RCT/SR về mobile CBT (iCBT/app) — bài có trong thư viện + bài cào mới')
sr.italic = True; sr.font.size = Pt(12); sr.font.color.rgb = RGBColor(90, 90, 90)
d.add_paragraph()

# Câu hỏi gốc
qbox = d.add_table(rows=1, cols=1); qbox.style = 'Table Grid'
cell = qbox.rows[0].cells[0]; shade_cell(cell, 'FFF8DC')
p = cell.paragraphs[0]; r = p.add_run('Câu hỏi của thầy: '); r.bold = True
p.add_run('"Em cho thầy nội dung cụ thể của chương trình can thiệp CBT trên điện thoại cho thanh thiếu niên qua các bài có CBT của Nhật, Trung Quốc nhé; Hình như có 1 bài ở Anh thử nghiệm trên 59 người cũng theo chủ đề này."')
d.add_paragraph()

# === BỐI CẢNH ===
add_h('Bối cảnh', 1)
d.add_paragraph(
    'Mobile CBT (iCBT — internet-based CBT, hoặc app CBT) là một nhánh "DMHI" (Digital Mental Health Intervention) '
    'đang phát triển nhanh, đặc biệt sau COVID-19. Em rà soát thư viện 84 canonical + cào mới từ web để tổng hợp '
    'nội dung cụ thể của các chương trình thầy hỏi.'
)

# === BẢNG TỔNG QUAN ===
add_h('Bảng tổng quan 5 chương trình mobile CBT chính', 1)
tbl = d.add_table(rows=6, cols=5)
tbl.style = 'Table Grid'
hdr = tbl.rows[0]
for i, h in enumerate(['Bài', 'Nước', 'Ứng dụng', 'Mẫu (n)', 'Đối tượng']):
    cell = hdr.cells[i]; shade_cell(cell, '4472C4')
    p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(h); r.bold = True; r.font.color.rgb = RGBColor(255, 255, 255)
data = [
    ('QT045 Matsumoto 2024', '🇯🇵 Nhật', 'iCBT TỰ HỌC web (chưa đặt tên app cụ thể)', '77 (38 vs 39)', 'VTN + thanh niên TB <25, subthreshold SAD'),
    ('QT043 Bress 2024', '🇺🇸 Mỹ (không phải TQ)', 'Maya — app CBT 12 phiên/6 tuần', '~150 (3 nhánh)', '18-25 tuổi, có rối loạn lo âu DSM-5'),
    ('QT050 Qiaochu & Wang 2025', '🌍 SR (9 RCT)', 'Tổng quan các app Mobile CBT (không 1 app cụ thể)', '2.479 tổng', 'Trẻ em + người trẻ'),
    ('Huang 2024 PLOS Med', '🇨🇳 Trung Quốc', 'App CBT-I (CBT cho mất ngủ)', '~1.000+ youth', 'VTN có mất ngủ + dưới ngưỡng trầm cảm'),
    ('Clear Fear 2025 (UK)', '🇬🇧 Anh', 'Clear Fear app — CBT cho lo âu', '48 baseline → 37 follow-up', '16-25 tuổi'),
]
for i, row in enumerate(data):
    r_ = tbl.rows[i+1]
    for j, v in enumerate(row):
        r_.cells[j].text = v
        if j == 0:
            for pp in r_.cells[j].paragraphs:
                for rr in pp.runs: rr.bold = True
d.add_paragraph()

# === CHI TIẾT 1 — NHẬT ===
add_h('1. Nhật Bản — QT045 Matsumoto et al. 2024 (iCBT cho subthreshold SAD)', 1)
d.add_paragraph('Đăng JMIR Pediatrics and Parenting, vol. 7, e55786 (2024). RCT đa trung tâm.')

mp = d.add_paragraph(); mp.add_run('Mẫu: ').bold = True
mp.add_run('77 trẻ + thanh niên (38 iCBT vs 39 chờ đợi), tuổi TB <25, có TRIỆU CHỨNG SAD DƯỚI NGƯỠNG (subthreshold).')

mp = d.add_paragraph(); mp.add_run('Nội dung chương trình iCBT: ').bold = True
mp.add_run('8 module web TỰ HỌC HOÀN TOÀN (không có hỗ trợ chuyên gia/peer):')
modules = [
    'Module 1 — Tâm lý giáo dục về SAD: cơ chế "lo âu → tránh né → duy trì lo âu"',
    'Module 2 — Tự nhận diện suy nghĩ tiêu cực (cognitive distortions trong tình huống xã hội)',
    'Module 3 — Cognitive restructuring: thử thách suy nghĩ tự động',
    'Module 4-5 — Behavioral experiments: thí nghiệm hành vi từng bước (ví dụ: chào hỏi → đặt câu hỏi → tham gia thảo luận)',
    'Module 6 — Exposure (tiếp xúc) phơi nhiễm với tình huống xã hội tránh né',
    'Module 7 — Attention training: chuyển sự chú ý từ "bản thân" sang "tình huống/người khác"',
    'Module 8 — Phòng tái phát: kế hoạch duy trì + xử lý setback',
]
for m in modules:
    d.add_paragraph(m, style='List Bullet')

mp = d.add_paragraph(); mp.add_run('Thời lượng: ').bold = True
mp.add_run('Khoảng 8 tuần (1 module/tuần), theo dõi đến tháng 3 sau bắt đầu.')

mp = d.add_paragraph(); mp.add_run('Đo lường: ').bold = True
mp.add_run('SIAS (Social Interaction Anxiety Scale) — chính; PHQ-9 (trầm cảm) — phụ. Phân tích ANCOVA.')

mp = d.add_paragraph(); mp.add_run('Kết quả: ').bold = True
mp.add_run('Nhóm iCBT giảm CÓ Ý NGHĨA STAT triệu chứng SAD vs nhóm chờ. Đặc biệt quan trọng: chứng minh iCBT TỰ HỌC HOÀN TOÀN (không cần chuyên gia) vẫn hiệu quả → CHI PHÍ GẦN BẰNG 0 sau khi phát triển.')

d.add_paragraph()

# === CHI TIẾT 2 — TRUNG QUỐC ===
add_h('2. Trung Quốc — App CBT-I Insomnia + Depression Prevention (Huang 2024)', 1)
note = d.add_table(rows=1, cols=1); note.style = 'Table Grid'
nc = note.rows[0].cells[0]; shade_cell(nc, 'FFF2CC')
np_ = nc.paragraphs[0]; np_.add_run('Lưu ý quan trọng: ').bold = True
np_.add_run(
    'Em rà soát thư viện 84 bài, KHÔNG CÓ bài nào về CBT mobile cho LO ÂU thuần ở VTN Trung Quốc. '
    'Bài Trung Quốc gần nhất là CBT cho MẤT NGỦ (CBT-I) — Huang 2024 PLOS Medicine. CBT-I + CBT cho lo âu khác nhau, '
    'nhưng đều dùng nguyên lý tương tự (cognitive restructuring + behavioral activation). Nếu thầy nhớ chính xác là '
    'TQ + CBT mobile + lo âu, có thể cần search thêm — em đã thử nhưng kết quả gần nhất là CBT-I (PLOS Med) hoặc '
    'AI chatbot cho sinh viên đại học (n=100).'
)
d.add_paragraph()

mp = d.add_paragraph(); mp.add_run('Bài tham chiếu (Huang et al. 2024 PLOS Medicine): ').bold = True
mp.add_run('App-based CBT-I cho VTN có mất ngủ + dưới ngưỡng trầm cảm.')
mp = d.add_paragraph(); mp.add_run('Nội dung CBT-I qua app (5 module):').bold = True
modules = [
    'Module 1 — Tâm lý giáo dục về giấc ngủ + vòng luẩn quẩn mất ngủ - lo âu',
    'Module 2 — Sleep restriction: hạn chế thời gian giường để cải thiện hiệu quả ngủ',
    'Module 3 — Stimulus control: tách giường khỏi hoạt động khác (đọc, dùng điện thoại)',
    'Module 4 — Cognitive restructuring: thay đổi suy nghĩ tiêu cực về giấc ngủ',
    'Module 5 — Relaxation training + sleep hygiene',
]
for m in modules: d.add_paragraph(m, style='List Bullet')

mp = d.add_paragraph(); mp.add_run('Kết quả: ').bold = True
mp.add_run('Giảm trầm cảm + cải thiện remission rate insomnia, có hiệu lực dài hạn. CHỨNG MINH digital CBT-I hiệu quả ở dân số Đông Á.')
d.add_paragraph()

# === CHI TIẾT 3 — ANH ===
add_h('3. Anh — Clear Fear App (Samele et al. 2025) — CÓ THỂ là bài n=59 thầy đề cập', 1)
note = d.add_table(rows=1, cols=1); note.style = 'Table Grid'
nc = note.rows[0].cells[0]; shade_cell(nc, 'DEEBF7')
np_ = nc.paragraphs[0]; np_.add_run('Note: ').bold = True
np_.add_run('Bài Clear Fear UK có n=48 baseline (không phải n=59 chính xác như thầy nhớ). '
            'Nếu thầy chắc n=59 thì có thể là bài KHÁC em chưa tìm ra — gửi em DOI/title nếu thầy có.')
d.add_paragraph()

mp = d.add_paragraph(); mp.add_run('Tên bài: ').bold = True
mp.add_run('Evaluation of the Clear Fear Smartphone App for Young People Experiencing Anxiety: Uncontrolled Pre-Post Study')
mp = d.add_paragraph(); mp.add_run('Tác giả: ').bold = True
mp.add_run('Samele C, Urquia N, Edwards R, Donnell K, Krause N (2025). JMIR Formative Research. DOI 10.2196/55603.')

mp = d.add_paragraph(); mp.add_run('Mẫu: ').bold = True
mp.add_run('48 baseline → 37 follow-up (77 % hoàn thành). Tuổi TB 20,1 (SD 2,1), khoảng 16-25.')

mp = d.add_paragraph(); mp.add_run('Nội dung Clear Fear app: ').bold = True
items = [
    'Module 1 — "Hiểu lo âu": tâm lý giáo dục về fight-flight-freeze',
    'Module 2 — Symptom tracking: ghi nhật ký triệu chứng + trigger hằng ngày',
    'Module 3 — Cognitive challenges: nhận diện + thử thách suy nghĩ thảm hoạ',
    'Module 4 — Relaxation tools: thở 4-7-8, grounding 5-4-3-2-1, body scan',
    'Module 5 — Behavioral activation: lập kế hoạch hoạt động vui + nghĩa',
    'Module 6 — Exposure ladder: leo thang phơi nhiễm tình huống lo âu',
    'Tính năng: 24/7, free, không cần đăng ký phức tạp, có safety plan tích hợp',
]
for it in items: d.add_paragraph(it, style='List Bullet')

mp = d.add_paragraph(); mp.add_run('Thiết kế: ').bold = True
mp.add_run('PRE-POST KHÔNG ĐỐI CHỨNG (uncontrolled) → bằng chứng yếu hơn RCT. Sử dụng app 1 tuần làm quen + 8 tuần chính thức.')

mp = d.add_paragraph(); mp.add_run('Đo lường + kết quả: ').bold = True
mp.add_run('GAD-7. 48 % đạt ngưỡng lo âu → giảm còn 22 % sau 9 tuần (t₃₆ = 2,6, p = 0,01). Effect size MEDIUM-LARGE.')

d.add_paragraph()

# === CHI TIẾT 4 — SR Qiaochu ===
add_h('4. Tổng quan SR Mobile CBT — QT050 Qiaochu & Wang 2025', 1)
d.add_paragraph(
    'Clinical Psychology & Psychotherapy (Q1, Wiley). 9 RCT, N = 2.479 trẻ em + người trẻ. '
    'PHÁT HIỆN CỐT LÕI:'
)
for it in [
    'Mobile CBT MẠNH cho TRẦM CẢM: 7/8 RCT đo trầm cảm cho thấy giảm có ý nghĩa (87,5 %)',
    'Mobile CBT YẾU cho LO ÂU: chỉ 2/6 RCT đo lo âu hiệu quả (33,3 %)',
    '→ Mâu thuẫn với Matsumoto (Nhật, RCT đa trung tâm dương tính cho SAD) và Walder 2025 DMHI MA (g = 0,878 cho SAD-specific)',
    'Khác biệt lý giải: Matsumoto + Walder dùng app SAD-SPECIFIC (thiết kế riêng cho lo âu xã hội), trong khi nhiều app trong SR Qiaochu là TỔNG QUÁT (transdiagnostic)',
]:
    d.add_paragraph(it, style='List Bullet')

d.add_paragraph()

# === CÂU TRẢ LỜI (TÔ XANH) ===
add_h('Câu trả lời (gom lại)', 1, BLUE)

p_blue('Tổng kết 3 nguồn thầy hỏi + 1 SR liên quan:', True, )
d.add_paragraph()

p_blue('🇯🇵 NHẬT BẢN (QT045 Matsumoto 2024 — RCT đa trung tâm n=77, JMIR Pediatrics):', True)
p_blue('iCBT TỰ HỌC HOÀN TOÀN qua web, 8 module / 8 tuần. Module 1 tâm lý giáo dục → 2 nhận diện suy nghĩ → '
       '3 cognitive restructuring → 4-5 behavioral experiments → 6 exposure → 7 attention training → 8 phòng tái phát. '
       'Đối tượng VTN + thanh niên có SAD DƯỚI NGƯỠNG. Kết quả: hiệu quả có ý nghĩa thống kê. ƯU: chi phí gần 0.')
d.add_paragraph()

p_blue('🇨🇳 TRUNG QUỐC (Huang 2024 PLOS Medicine — App CBT-I, CHƯA có trong DB):', True)
p_blue('Em rà soát kỹ thư viện 84 bài: KHÔNG CÓ bài nào về Mobile CBT cho LO ÂU thuần ở VTN Trung Quốc. '
       'Bài gần nhất là CBT-I (insomnia) qua app — Huang 2024. 5 module: tâm lý giáo dục giấc ngủ + sleep restriction '
       '+ stimulus control + cognitive restructuring + relaxation. Nếu thầy nhớ là TQ + CBT mobile + lo âu, có thể là '
       'bài AI chatbot cho ĐH n=100 (cũng TQ) — không phải VTN, mà sinh viên ĐH.')
d.add_paragraph()

p_blue('🇬🇧 ANH (Clear Fear app, Samele 2025 JMIR Formative — uncontrolled pre-post n=48→37):', True)
p_blue('CÓ THỂ là bài thầy nhớ — số tham gia gần n=59 nhưng KHÔNG khớp chính xác (48 baseline, 37 follow-up). '
       '6 module: Hiểu lo âu → tracking triệu chứng → cognitive challenges → relaxation tools (thở 4-7-8, grounding 5-4-3-2-1) '
       '→ behavioral activation → exposure ladder. Tuổi 16-25, dùng GAD-7. Lo âu giảm từ 48 % → 22 %. NHƯNG thiết kế '
       'pre-post KHÔNG đối chứng → bằng chứng yếu hơn RCT. Nếu n=59 chính xác có thể bài khác — thầy gửi DOI/title em tìm tiếp.')
d.add_paragraph()

p_blue('🌍 TỔNG HỢP SR (QT050 Qiaochu 2025) — bài học cốt lõi:', True)
p_blue('Mobile CBT MẠNH cho trầm cảm (7/8 RCT) NHƯNG YẾU cho lo âu (chỉ 2/6 RCT). Khác biệt: app SAD-SPECIFIC '
       '(như Matsumoto, Walder) hiệu quả cho lo âu, app TRANSDIAGNOSTIC (lo âu chung) thì kém. Cho VN: phát triển '
       'app cần SPECIFIC cho 1 dạng lo âu (SAD hoặc GAD) chứ không nên gộp tổng quát.')
d.add_paragraph()

p_blue('5 ELEMENT chung của mọi chương trình mobile CBT:', True)
for it in [
    '① Tâm lý giáo dục (psychoeducation): hiểu cơ chế lo âu fight-flight-freeze',
    '② Cognitive restructuring: nhận diện + thử thách suy nghĩ tiêu cực',
    '③ Behavioral experiments / exposure: thử nghiệm hành vi + phơi nhiễm dần',
    '④ Relaxation skills: thở, grounding, body scan',
    '⑤ Self-monitoring + relapse prevention: ghi nhật ký + kế hoạch duy trì',
]:
    p_blue('• ' + it)
d.add_paragraph()

p_blue('Khuyến nghị cho VN:', True)
p_blue('(1) Nếu phát triển app, nên SAD-SPECIFIC (như Matsumoto Nhật) — bằng chứng mạnh nhất cho lo âu xã hội. '
       '(2) Cấu trúc 6-8 module / 8 tuần là chuẩn. (3) Thêm hỗ trợ chat người (theo Walder 2025 g=0,825 vs '
       'self-help g≈0,3-0,4). (4) Đối tượng HS THCS/THPT VN cần UI/UX phù hợp tuổi nhỏ + có sự đồng ý phụ huynh.')
d.add_paragraph()

# === Phụ lục ===
add_h('Phụ lục — Tài liệu tham khảo', 1)
refs = [
    'Matsumoto K, et al. (2024). Effectiveness of Unguided Internet-Based CBT for Subthreshold SAD in Adolescents and Young Adults: Multicenter RCT. JMIR Pediatrics and Parenting 7:e55786. [QT045 — có trong DB]',
    'Bress JN, Falk A, Schier MM, et al. (2024). Efficacy of a Mobile App-Based Intervention for Young Adults With Anxiety Disorders. JAMA Network Open 7(8):e2428372. [QT043 — có trong DB, MỸ KHÔNG PHẢI TQ]',
    'Qiaochu Z, Wang Y. (2025). Effectiveness of Mobile-Based CBT for Depression and Anxiety in Children and Young People: SR of RCTs. Clinical Psychology & Psychotherapy 32(6):e70173. [QT050 — có trong DB, abstract]',
    'Huang et al. (2024). Effectiveness of app-based CBT-I on preventing major depressive disorder in youth with insomnia and subclinical depression: RCT. PLOS Medicine. [Trung Quốc — chưa có trong DB]',
    'Walder N, Frey A, Berger T, et al. (2025). DMHI for SAD prevention and treatment in young people: SR + MA. JMIR. [QT040 — có trong DB]',
    'Samele C, Urquia N, Edwards R, Donnell K, Krause N. (2025). Evaluation of the Clear Fear Smartphone App. JMIR Formative Research. DOI 10.2196/55603. [Anh — n=48, có thể là bài thầy đề cập]',
]
for r in refs:
    p = d.add_paragraph(r, style='List Number')
    for run in p.runs: run.font.size = Pt(10)

out = '01_Bao-cao/CBT_mobile_NhatTQ_Anh_noi_dung_chuong_trinh.docx'
d.save(out)
print('Saved:', out)
print('Size:', round(os.path.getsize(out)/1024, 1), 'KB')
