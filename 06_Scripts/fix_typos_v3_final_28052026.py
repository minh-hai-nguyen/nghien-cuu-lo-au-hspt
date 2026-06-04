# -*- coding: utf-8 -*-
"""Sua not 10 double-space + 4 space-before-punct con sot.
28/05/2026."""
import os, sys, io, re
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

d = Document(FILE)
stats = {'double_space': 0, 'space_before_punct': 0}

# =========================================================
# 1. Double space - apply across full doc
# (My earlier run only worked at run-level. Some doubles might span runs.)
# =========================================================
print("[1] Sua double-space across runs...")
for i, p in enumerate(d.paragraphs):
    if '  ' not in p.text:
        continue
    # Build full text, fix, then redistribute
    full = p.text
    fixed_full = re.sub(r'  +', ' ', full)
    if fixed_full == full:
        continue
    # Apply at run-level (assume formatting consistent)
    for run in p.runs:
        if '  ' in run.text:
            run.text = re.sub(r'  +', ' ', run.text)
            stats['double_space'] += 1
    # If still has double-space (spans runs), join+re-split won't preserve formatting.
    # We accept the run-level fix.


# =========================================================
# 2. Specific space-before-punct fixes (text might span runs)
# =========================================================
print("\n[2] Sua space-before-punct con sot...")
patterns_paragraphs = [
    (502, [('lan tỏa ,', 'lan tỏa,')]),
    (1217, [('ISSN : ', 'ISSN: ')]),
    (1316, [('DOI : ', 'DOI: ')]),
    (1325, [('DOI : ', 'DOI: ')]),
]
for idx, repls in patterns_paragraphs:
    if idx >= len(d.paragraphs): continue
    p = d.paragraphs[idx]
    for run in p.runs:
        for old, new in repls:
            if old in run.text:
                run.text = run.text.replace(old, new)
                stats['space_before_punct'] += 1
                print(f"  Para {idx}: '{old}' -> '{new}' (in run)")
    # Also try cross-run: rebuild
    full = p.text
    fixed = full
    for old, new in repls:
        fixed = fixed.replace(old, new)
    if fixed != full:
        # Cross-run pattern, need to fix at full-text level
        # Reduce to single run with original first-run formatting
        ## Actually skip — could break formatting.
        pass


# =========================================================
# 3. Aggressively find " ," " ;" " :" patterns across full doc
# (Vietnamese: never space before , ; : ! ?)
# =========================================================
print("\n[3] Quet toan bo doc tim ' ,' ' ;' ' :' (across runs)...")
fixed_cross_run = 0
for i, p in enumerate(d.paragraphs):
    if not p.runs: continue
    full = p.text
    # Skip URL-heavy paras (TLTK with DOIs)
    if 'http' in full.lower() or 'doi' in full.lower() or 'pmid' in full.lower():
        # Only fix specific patterns
        targets = [' :', ' ;']
    else:
        targets = [' ,', ' ;', ' :']
    has_pattern = False
    for t in targets:
        # Pattern: letter + space + punctuation
        if re.search(rf'\w{re.escape(t)}', full):
            has_pattern = True
            break
    if not has_pattern: continue
    # Try at run level first
    for run in p.runs:
        for t in targets:
            # Replace 'X ,' -> 'X,'
            new_text = re.sub(rf'(\w)\s+([{re.escape(t.strip())}])(?!\d)', r'\1\2', run.text)
            if new_text != run.text:
                run.text = new_text
                fixed_cross_run += 1
print(f"  Fixed (at run level): {fixed_cross_run}")


d.save(FILE)
print(f"\nSaved. Stats: {stats}, cross_run: {fixed_cross_run}")
