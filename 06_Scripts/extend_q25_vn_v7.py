# -*- coding: utf-8 -*-
"""
Q2.5 VN v6 -> v7: extend nhẹ — bổ sung 9 đoạn bản EN đang có nhưng bản VN thiếu,
đưa hai bản song song. Toàn bộ là DỊCH nội dung đã có + đã xác minh ở bản EN v6
(không thêm dữ liệu mới). Bản EN giữ nguyên v6.
9 đoạn: [12] bối cảnh GDPT 2018; [23] lực thống kê; [24] đối chiếu chuẩn quốc tế;
[40] hàm ý sàng lọc; [41] suy ngẫm đặc thù phân nhóm; [42] đối chiếu Đông Á;
[55] liên hệ 2 bài tổng quan đồng hành; [56] xung đột lợi ích & tài trợ;
[57] cơ sở bằng chứng theo phân nhóm.
"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document
from docx.shared import Pt

IN = "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v6.docx"
OUT = "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v7.docx"

# {heading text: [đoạn cần chèn TRƯỚC heading đó, theo thứ tự]}
INSERTS = {
 "2. Phương pháp nghiên cứu": [
  "Bối cảnh Việt Nam của nghiên cứu này được định hình bởi cải cách giáo dục sau "
  "năm 2018, với Chương trình giáo dục phổ thông 2018 chú trọng rõ rệt đến dạy học "
  "theo năng lực và tích hợp nội dung kỹ năng sống cùng sức khỏe tâm thần vào "
  "chương trình chính khóa. Hoạt động trải nghiệm, hướng nghiệp — một chương trình "
  "hằng tuần trong chương trình này — tạo ra một kênh triển khai tự nhiên cho can "
  "thiệp sức khỏe tâm thần dựa trên trường học, song hiện chưa có các thành phần "
  "CBT có cấu trúc như trong các thử nghiệm phòng ngừa dựa trên trường học của "
  "quốc tế. Các phát hiện của nghiên cứu này hướng tới việc cung cấp cơ sở cho nội "
  "dung tương lai của khung giờ chương trình đó.",
 ],
 "3. Kết quả nghiên cứu": [
  "Cân nhắc về lực thống kê và độ nhạy. Phân tích lực thống kê (power) dựa trên "
  "mẫu phân tích 1.352 cho thấy nghiên cứu có lực trên 99% để phát hiện các hệ số "
  "đường dẫn chuẩn hóa cỡ nhỏ đến trung bình (β ≥ 0,10) ở mức α = 0,05 trong các "
  "mô hình đường dẫn đơn, và trên 95% để phát hiện các hiệu ứng như vậy trong mô "
  "hình đa đường dẫn tích hợp. Ngưỡng phát hiện cho hiệu ứng tương tác trong phân "
  "tích điều tiết được ước tính ở β ≥ 0,08. Các con số này cho thấy nghiên cứu có "
  "lực thống kê tốt cho các câu hỏi nghiên cứu chính; do đó các phát hiện không có "
  "ý nghĩa (đáng chú ý là đường dẫn hỗ trợ cha mẹ bằng không lên lo âu chia ly) "
  "khó có khả năng phản ánh sai lầm loại II.",
  "Đối chiếu với các chuẩn phân tích tổng hợp quốc tế. Để đặt độ lớn báo cáo ở "
  "trên vào bối cảnh, chúng tôi đối chiếu từng đường dẫn với các ước lượng phân "
  "tích tổng hợp quốc tế đã công bố khi có sẵn. Đường dẫn β = 0,510 từ áp lực học "
  "tập đến lo âu tổng quát trong mẫu này vượt mức tương quan trung bình ghi nhận "
  "trong y văn quốc tế về stress học đường (thường r = 0,30–0,45), phù hợp với "
  "giả thuyết rằng các hệ thống giáo dục Đông Á nặng về thi cử khuếch đại liên hệ "
  "stress học tập – lo âu. Đường dẫn β = 0,376 từ bắt nạt học đường đến lo âu chia "
  "ly không thể đặt cùng một hệ quy chiếu với tỷ số chênh OR = 1,77 mà Moore và "
  "cộng sự báo cáo cho lo âu nói chung, vì hệ số hồi quy chuẩn hóa và tỷ số chênh "
  "không so sánh trực tiếp được; tuy vậy, phát hiện đặc thù theo phân nhóm ở đây "
  "tinh chỉnh bằng chứng phân tích tổng hợp đó bằng cách khu trú liên hệ bắt nạt – "
  "lo âu vào phân nhóm lo âu chia ly.",
 ],
 "4. Bàn luận": [
  "Hàm ý cho sàng lọc thường quy tại trường học. Các phát hiện đường dẫn có hàm ý "
  "trực tiếp cho việc thiết kế quy trình sàng lọc thường quy tại các trường trung "
  "học cơ sở Việt Nam. Một bộ công cụ sàng lọc nhận biết phân nhóm — phân biệt lo "
  "âu tổng quát, lo âu xã hội và lo âu chia ly thay vì gộp chung — sẽ cho phép "
  "phân loại học sinh vào các hướng can thiệp riêng theo đường dẫn. Ví dụ, học "
  "sinh dương tính với lo âu chia ly sẽ hưởng lợi từ việc đánh giá và can thiệp "
  "chống bắt nạt, trong khi học sinh dương tính với lo âu tổng quát hoặc lo âu xã "
  "hội sẽ hưởng lợi từ quản lý áp lực học tập và các quy trình tăng cường lòng tự "
  "trọng. Sàng lọc nhận biết phân nhóm như vậy là cách chuyển dịch thực tiễn các "
  "phát hiện cấp đường dẫn của nghiên cứu vào quy trình thường nhật của dịch vụ "
  "sức khỏe tâm thần học đường.",
  "Suy ngẫm về các phát hiện đặc thù phân nhóm. Mẫu hình đường dẫn khác biệt giữa "
  "ba phân nhóm lo âu là một trong những đóng góp mới về mặt lý thuyết của nghiên "
  "cứu này, và việc lặp lại nó ở các bối cảnh tập thể châu Á khác sẽ là một ưu "
  "tiên nghiên cứu cao. Mẫu hình này nhất quán nội tại: bắt nạt học đường — một "
  "yếu tố căng thẳng khu trú ở trường — dự báo lo âu chia ly mạnh hơn hẳn; áp lực "
  "học tập — một yếu tố căng thẳng nhận thức – đánh giá khu trú ở trường — dự báo "
  "lo âu tổng quát và lo âu xã hội mạnh hơn hẳn; hỗ trợ cha mẹ — một yếu tố bảo vệ "
  "khu trú ở gia đình — bảo vệ chọn lọc khỏi lo âu xã hội; và lòng tự trọng — một "
  "nguồn lực nhận thức – đánh giá nội tại — song hành với áp lực học tập xuyên các "
  "phân nhóm nhận thức – đánh giá. Mẫu hình hội tụ này củng cố độ tin cậy diễn "
  "giải của từng phát hiện riêng lẻ.",
  "Đối chiếu với khả năng lặp lại ở Đông Á. Mặc dù các nghiên cứu so sánh trực "
  "tiếp còn hiếm, mẫu hình quan sát được ở đây có thể được tam giác hóa một cách "
  "thận trọng với các phát hiện đã công bố từ các hệ thống giáo dục di sản Khổng "
  "giáo lân cận. Các nghiên cứu từ Trung Quốc đại lục và Hồng Kông đã ghi nhận "
  "đường dẫn stress học tập – lo âu lớn tương tự và mức gắn bó cha mẹ nền cao "
  "tương tự, song so sánh trực tiếp theo từng đường dẫn – phân nhóm thì chưa có "
  "trong y văn quốc tế đã công bố. Nghiên cứu này do đó vừa là một đóng góp cho "
  "bối cảnh Việt Nam vừa là một khuôn mẫu phương pháp cho nghiên cứu SEM so sánh "
  "Đông Á trong tương lai về đường dẫn lo âu vị thành niên. Chúng tôi dự đoán rằng "
  "công trình so sánh như vậy sẽ làm lộ ra cả mẫu hình chung lẫn khác biệt giữa "
  "các hệ thống giáo dục di sản Khổng giáo, với đặc điểm chung gồm đường dẫn áp "
  "lực học tập mạnh và đặc điểm khác biệt gồm các yếu tố điều tiết văn hóa đặc thù "
  "của hỗ trợ cha mẹ và lòng tự trọng.",
 ],
 "5. Kết luận": [
  "Liên hệ với hai bài tổng quan đồng hành của nhóm. Hai bài tổng quan tường "
  "thuật của nhóm chúng tôi gần đây đã tổng hợp bằng chứng Việt Nam và quốc tế đã "
  "công bố về các yếu tố nguy cơ và về nghiên cứu can thiệp cho rối loạn lo âu ở "
  "học sinh trung học cơ sở Việt Nam (Authors, 2026a; Authors, 2026b). Nghiên cứu "
  "thực nghiệm hiện tại mở rộng hai bài tổng quan đó bằng cách cung cấp ước lượng "
  "cấp đường dẫn từ một mẫu Việt Nam lớn, tách hiệu ứng theo các phân nhóm lo âu, "
  "và tích hợp cả yếu tố nguy cơ lẫn yếu tố bảo vệ trong một mô hình cấu trúc "
  "tuyến tính duy nhất. Tính bổ trợ của ba đóng góp — hai bài tổng quan và một "
  "nghiên cứu thực nghiệm — nhằm hỗ trợ phát triển thiết kế can thiệp nhận biết "
  "phân nhóm trong sức khỏe tâm thần học đường Việt Nam.",
  "Xung đột lợi ích và tài trợ. Các tác giả tuyên bố không có xung đột lợi ích. "
  "Việc thu thập dữ liệu được thực hiện như một phần của đào tạo tiến sĩ và không "
  "nhận tài trợ thương mại bên ngoài. Bộ dữ liệu đã ẩn danh sẽ được cung cấp cho "
  "các nhà nghiên cứu đủ điều kiện theo yêu cầu hợp lý, với sự chấp thuận của hội "
  "đồng đạo đức. Chúng tôi trân trọng cảm ơn các trường, học sinh, phụ huynh và "
  "giáo viên đã hợp tác để nghiên cứu này được thực hiện.",
  "Cơ sở bằng chứng theo phân nhóm. Tổng hợp lại, các phát hiện đường dẫn, phân "
  "tích trung gian và các kiểm tra độ nhạy báo cáo ở trên đều nhất quán chỉ tới "
  "một cơ sở bằng chứng được thiết kế theo phân nhóm, trong đó các phân nhóm lo âu "
  "khác nhau có hồ sơ yếu tố nguy cơ chồng lấp một phần nhưng khác biệt đáng kể. "
  "Chuyển dịch cơ sở bằng chứng này vào thực tiễn sẽ đòi hỏi nỗ lực phối hợp giữa "
  "sàng lọc tại trường học, đào tạo giáo viên, can thiệp ở cấp gia đình và việc "
  "tiếp tục phát triển các công cụ đánh giá đã được kiểm định văn hóa cho vị thành "
  "niên Việt Nam.",
 ],
}

d = Document(IN)


def find_heading(doc, text):
    for p in doc.paragraphs:
        if p.text.strip() == text:
            return p
    return None


total = 0
for heading_text, paras in INSERTS.items():
    h = find_heading(d, heading_text)
    if h is None:
        print(f"!! KHÔNG tìm thấy heading: {heading_text}")
        continue
    for t in paras:
        np = h.insert_paragraph_before()
        r = np.add_run(t)
        r.font.name = "Times New Roman"
        r.font.size = Pt(12)
        total += 1
    print(f"  Chèn {len(paras)} đoạn trước '{heading_text}'")

print(f"Tổng đoạn đã chèn: {total}")
d.core_properties.author = ""
d.core_properties.last_modified_by = ""
d.save(OUT)
print(f"Đã lưu: {OUT}  | tổng paragraphs: {len(d.paragraphs)}")
