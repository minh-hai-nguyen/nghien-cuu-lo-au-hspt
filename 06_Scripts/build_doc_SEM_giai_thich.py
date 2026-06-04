"""DOC: Giai thich SEM (Structural Equation Modeling) — beta chuan hoa, fit indices.

ALL FACTS VERIFIED:
- Hu & Bentler 1999 cutoffs: CFI/TLI ≥ 0,95; RMSEA ≤ 0,06; SRMR ≤ 0,08 - WebSearch verified
- DOI: 10.1080/10705519909540118
- beta standardized = beta unstandardized × (SD_X / SD_Y) - canonical
- Cohen 1988 effect size for SEM: |β| ≥ 0,1 nho; ≥ 0,3 trung binh; ≥ 0,5 lon
- Chuong 3 luan an: beta ALHT 0,510/0,490; beta TTr -0,455/-0,415; CFI > 0,99 - verified
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/SEM_giai_thich_beta_chuan_hoa_fit_indices.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(2.5); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:12}.get(level, 12))
    r.font.color.rgb = color

def para(text, size=12, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(11)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(11)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('SEM (STRUCTURAL EQUATION MODELING) LÀ GÌ?\nβ CHUẨN HÓA, FIT INDICES, VÀ ỨNG DỤNG\nCHO LUẬN ÁN LO ÂU HỌC SINH THCS\n— Giải thích chi tiết theo CHƯƠNG 3 luận án —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# 1. SEM là gì?
H('1. SEM là gì?', level=2, color=NAVY)
para(
    'SEM (Structural Equation Modeling — mô hình phương trình cấu '
    'trúc) là khung phân tích thống kê đa biến CỘNG HỢP giữa hai '
    'kỹ thuật cổ điển: (1) phân tích nhân tố khẳng định '
    '(Confirmatory Factor Analysis — CFA) cho biến tiềm ẩn '
    '(latent variables), và (2) hồi quy đường dẫn (path analysis) '
    'cho mối quan hệ giữa các biến. Nói cách khác, SEM cho phép '
    'kiểm định ĐỒNG THỜI nhiều mối quan hệ giữa các biến tiềm ẩn '
    'và biến quan sát.', indent=False
)
para('SEM khác gì với hồi quy bội (multiple regression)?', bold=True, indent=False)
add_table(
    ['Đặc điểm', 'Hồi quy bội cổ điển', 'SEM'],
    [
        ['Số biến phụ thuộc', '1', 'Có thể nhiều (hệ thống đường dẫn)'],
        ['Biến tiềm ẩn', 'Không hỗ trợ', 'Có (latent variables)'],
        ['Sai số đo (measurement error)', 'Bỏ qua — giả định đo chính xác', 'Mô hình hóa rõ ràng'],
        ['Trung gian + điều tiết', 'Cần phân tích bước', 'Trực tiếp trong một mô hình'],
        ['Đánh giá mô hình', 'R², F-test, residuals', 'CFI, TLI, RMSEA, SRMR, χ²/df + R²'],
        ['Cỡ mẫu khuyến nghị', '~10–20 obs/biến', 'Tối thiểu n = 200; tốt n ≥ 400'],
    ]
)
para('')
para(
    'Khi nào DÙNG SEM thay vì hồi quy bội? Bốn trường hợp:'
)
para(
    '1. Có biến tiềm ẩn — đo bằng nhiều mục (ví dụ "lo âu" đo bằng '
    '7 mục GAD-7).\n'
    '2. Cần kiểm định trung gian (mediation) hoặc điều tiết '
    '(moderation) trong cùng mô hình.\n'
    '3. Có nhiều biến phụ thuộc đồng thời (ví dụ ALHT → RLLATQ + '
    'RLLAXH + RLLAC).\n'
    '4. Cần kiểm định MÔ HÌNH lý thuyết tổng thể, không chỉ từng '
    'mối quan hệ riêng lẻ.', indent=False
)

# 2. β chuẩn hóa vs β chưa chuẩn hóa
H('2. β chuẩn hóa (standardized) vs β chưa chuẩn hóa (unstandardized)', level=2, color=NAVY)
para(
    'Trong SEM, mỗi đường ảnh hưởng (path) được biểu diễn bằng MỘT '
    'hệ số. Hệ số có thể ở hai dạng — mang ý nghĩa khác nhau.', indent=False
)
caption('Bảng 1. So sánh β chuẩn hóa và chưa chuẩn hóa')
add_table(
    ['Đặc điểm', 'β chưa chuẩn hóa (B)', 'β chuẩn hóa (β)'],
    [
        ['Đơn vị', 'Đơn vị gốc của X và Y', 'Không có đơn vị (unitless)'],
        ['Diễn giải', 'X tăng 1 đơn vị → Y tăng B đơn vị',
         'X tăng 1 SD → Y tăng β SD'],
        ['Khoảng giá trị', 'Không giới hạn',
         'Thường trong khoảng −1 đến 1 (mô hình đơn)'],
        ['So sánh giữa biến', 'KHÔNG so sánh được (đơn vị khác)',
         'SO SÁNH được (cùng đơn vị SD)'],
        ['Công thức quy đổi', 'B', 'β = B × (SD_X / SD_Y)'],
        ['Khi báo cáo', 'Cho dự đoán cụ thể', 'Cho so sánh cường độ tác động'],
    ]
)
para('')
para(
    'Trong báo cáo khoa học, ƯU TIÊN báo cáo β chuẩn hóa khi muốn '
    'SO SÁNH cường độ tác động giữa các biến. Khi cần dự đoán '
    'thực tế (ví dụ "tăng 10 phút screen time → tăng X điểm GAD-7"), '
    'báo cáo thêm β chưa chuẩn hóa.'
)
para('Ví dụ áp dụng cho chương 3 luận án:', bold=True, indent=False)
para(
    'Báo cáo "β = 0,510" của áp lực học tập với lo âu lan tỏa nghĩa '
    'là: khi áp lực học tập TĂNG 1 độ lệch chuẩn (SD), điểm lo âu '
    'lan tỏa TĂNG 0,510 SD — sau khi kiểm soát các biến đồng biến '
    'trong mô hình. Vì β = 0,510 > 0,4 và lớn hơn các β khác trong '
    'mô hình → áp lực học tập là yếu tố MẠNH NHẤT.'
)
para(
    'Khi so sánh với β = −0,455 của lòng tự trọng: |0,455| / |0,510| '
    '= 0,89 → tự trọng có cường độ NGANG BẰNG ~89% áp lực học tập '
    '— do hai hệ số ở cùng đơn vị SD, có thể so sánh trực tiếp.'
)

# 3. Fit indices — Hu & Bentler 1999
H('3. Fit indices — đánh giá mô hình SEM phù hợp với dữ liệu', level=2, color=NAVY)
para(
    'Sau khi ước lượng SEM, cần kiểm tra mô hình có PHÙ HỢP với dữ '
    'liệu thực không. Các chỉ số phổ biến — chuẩn vàng theo Hu và '
    'Bentler (1999) trong Structural Equation Modeling tập 6 số 1 '
    'trang 1–55 (DOI: 10.1080/10705519909540118):', indent=False
)
caption('Bảng 2. Chỉ số fit và ngưỡng theo Hu & Bentler (1999)')
add_table(
    ['Chỉ số', 'Tên đầy đủ', 'Ngưỡng "good fit"', 'Ngưỡng "acceptable"'],
    [
        ['CFI', 'Comparative Fit Index', '≥ 0,95', '≥ 0,90'],
        ['TLI', 'Tucker-Lewis Index (NNFI)', '≥ 0,95', '≥ 0,90'],
        ['RMSEA', 'Root Mean Square Error of Approximation', '≤ 0,06', '≤ 0,08'],
        ['SRMR', 'Standardized Root Mean Square Residual', '≤ 0,08', '≤ 0,10'],
        ['χ²/df', 'Chi-square / degrees of freedom', '< 3', '< 5'],
        ['χ² (Chi-square)', 'p-value', 'p > 0,05 (mô hình không bị bác)', 'Nhạy cảm với cỡ mẫu lớn'],
    ]
)
para('')
para(
    'Hu và Bentler (1999) khuyến nghị PHỐI HỢP HAI chỉ số (two-index '
    'presentation strategy) để giảm Type I và Type II errors: '
    'CFI ≥ 0,95 KẾT HỢP với SRMR ≤ 0,08 hoặc CFI ≥ 0,95 KẾT HỢP '
    'với RMSEA ≤ 0,06.'
)
para('Lưu ý quan trọng:', bold=True, indent=False)
para(
    '• χ² test rất NHẠY với cỡ mẫu — với n = 1.352 như chương 3 '
    'luận án, χ² gần như luôn có ý nghĩa thống kê (p < 0,05). '
    'KHÔNG nên dựa vào χ² đơn lẻ.\n'
    '• KTC 90% của RMSEA quan trọng — phải chứa ngưỡng 0,05 để '
    'kết luận "good fit".\n'
    '• RMSEA và TLI có XU HƯỚNG over-reject mô hình đúng khi cỡ '
    'mẫu nhỏ.', indent=False
)
para('Áp dụng cho chương 3 luận án:', bold=True, indent=False)
para(
    'Mô hình YTNC riêng (chỉ áp lực học tập + biện pháp đối phó '
    '+ ...) đạt CFI > 0,99 và KTC 90% RMSEA chứa ngưỡng 0,05 — '
    'mô hình PHÙ HỢP TỐT theo Hu & Bentler (1999). Mô hình tích '
    'hợp YTNC + YTBV cũng đạt fit indices tốt; R² = 0,598 thuộc '
    '"effect size lớn" theo Cohen (1988).'
)

# 4. Mediation và moderation
H('4. Trung gian (mediation) và điều tiết (moderation) trong SEM', level=2, color=NAVY)
para(
    'Hai khái niệm thường gây nhầm lẫn nhưng KHÁC NHAU về bản '
    'chất.', indent=False
)
caption('Bảng 3. So sánh trung gian và điều tiết')
add_table(
    ['Đặc điểm', 'Mediation (trung gian)', 'Moderation (điều tiết)'],
    [
        ['Câu hỏi', 'X tác động Y QUA M như thế nào?', 'X tác động Y MẠNH HƠN khi Z thế nào?'],
        ['Cấu trúc', 'X → M → Y (chuỗi)', 'X × Z → Y (tương tác)'],
        ['Kiểm định', 'Tác động gián tiếp (a × b)', 'Hệ số tương tác có ý nghĩa'],
        ['Phương pháp', 'Bootstrap KTC 95% cho a×b',
         'Hồi quy với biến tích X×Z'],
        ['Ví dụ luận án', 'Áp lực HT → giấc ngủ → lo âu',
         'Áp lực HT × giới tính → lo âu'],
    ]
)
para('')
para('Ví dụ áp dụng từ Zheng & Peng (2025) — QT041:', bold=True, indent=False)
para(
    'Mô hình mediation: Mạng xã hội (X) → Self-efficacy (M) → '
    'Lo âu (Y). Tổng tác động = 0,415 (r). Tác động gián tiếp '
    'qua self-efficacy = 63,13% × 0,415 ≈ 0,262. Tác động trực '
    'tiếp = 0,153 (β). Phù hợp với mô hình "phần lớn tác động '
    'của MXH lên lo âu là GIÁN TIẾP qua giảm self-efficacy".'
)

# 5. Cỡ mẫu cần cho SEM
H('5. Cỡ mẫu cần cho SEM', level=2, color=NAVY)
para(
    'Câu hỏi thường gặp: cỡ mẫu tối thiểu cho SEM là bao nhiêu?', indent=False
)
caption('Bảng 4. Khuyến nghị cỡ mẫu cho SEM (Kline, 2015)')
add_table(
    ['Cỡ mẫu', 'Đánh giá', 'Tình huống áp dụng'],
    [
        ['n < 100', 'Quá nhỏ', 'Không đủ cho SEM thông thường'],
        ['n = 100–200', 'Nhỏ', 'Mô hình đơn giản (≤ 3 latent var)'],
        ['n = 200–400', 'Trung bình', 'Mô hình trung bình (3–6 latent var)'],
        ['n ≥ 400', 'Lớn', 'Mô hình phức tạp với nhiều biến tiềm ẩn'],
        ['n ≥ 1.000', 'Rất lớn', 'Mô hình rất phức tạp + multi-group'],
    ]
)
para('')
para(
    'Quy tắc N:q ratio (Bentler & Chou, 1987): tỷ lệ cỡ mẫu trên '
    'số tham số ước lượng tối thiểu 5:1, lý tưởng 10:1. Ví dụ '
    'mô hình ước lượng 50 tham số → cần n ≥ 250 (tối thiểu) hoặc '
    'n ≥ 500 (lý tưởng).'
)
para(
    'Áp dụng cho chương 3 luận án: n = 1.352 — RẤT LỚN, vượt '
    'ngưỡng cho mô hình SEM phức tạp 6 biến tiềm ẩn (3 dạng '
    'RLLA + ALHT + TTr + BPĐP). Cỡ mẫu này cho phép kiểm định '
    'mô hình tích hợp YTNC + YTBV với độ tin cậy cao.'
)

# 6. Ứng dụng cho chương 3 luận án
H('6. Ứng dụng SEM trong chương 3 luận án', level=2, color=NAVY)
para(
    'Chương 3 luận án sử dụng SEM theo bốn lớp mô hình.', indent=False
)
para(
    'LỚP 1 — Mô hình ALHT → từng dạng RLLA (Bảng 3.24): '
    'β = 0,510 (lan tỏa), 0,490 (xã hội), p < 0,001. CFI > 0,99 '
    '— fit RẤT TỐT.'
)
para(
    'LỚP 2 — Mô hình TTr → từng dạng RLLA (Bảng 3.32): '
    'β = −0,455 (lan tỏa), −0,415 (xã hội), p < 0,001. Tỷ số '
    '|β TTr|/|β ALHT| = 0,85–0,89.'
)
para(
    'LỚP 3 — Mô hình BPĐP → từng dạng RLLA (Bảng 3.44–3.45): '
    'β = 0,749 (lan tỏa), 0,670 (xã hội). Cường độ CỰC LỚN '
    'nhưng RMSEA cao (0,080–0,204), CFI thấp (0,865–0,911) — '
    'mô hình fit CHƯA TỐT. Cần phân tách BPĐP theo '
    'Brief-COPE 14 nhân tố (Carver, 1997).'
)
para(
    'LỚP 4 — Mô hình tích hợp YTNC + YTBV (Bảng 3.37–3.38): '
    'β YTNC tổng = 0,669; β YTBV tổng = 0,220 (sau kiểm soát '
    'lẫn nhau). R² = 0,598 — effect size LỚN theo Cohen '
    '(1988). Mô hình tích hợp đạt fit indices tốt.'
)
para('Cảnh báo về diễn giải:', bold=True, indent=False)
para(
    'Mô hình YTNC riêng có |β| = 0,747 (CFI > 0,99), trong '
    'khi mô hình tích hợp YTNC + YTBV cho |β YTNC| = 0,669 '
    '— GIẢM nhẹ do YTBV cũng giải thích một phần phương sai. '
    'Đây là hiện tượng thông thường khi thêm biến cùng tác '
    'động lên outcome — KHÔNG phải lỗi mô hình.'
)

# 7. Software phổ biến
H('7. Software phổ biến cho SEM', level=2, color=NAVY)
caption('Bảng 5. Lựa chọn software cho SEM')
add_table(
    ['Software', 'Loại', 'Ưu', 'Nhược'],
    [
        ['lavaan (R)', 'Open-source', 'Miễn phí, syntax dễ, cộng đồng lớn', 'Đường cong học tập với người mới R'],
        ['Mplus', 'Thương mại', 'Mạnh nhất cho SEM phức tạp + multi-level', 'Đắt (~ $895 cho cá nhân)'],
        ['AMOS (IBM)', 'Thương mại', 'Giao diện đồ họa thân thiện', 'Hạn chế cho mô hình phức tạp'],
        ['SmartPLS', 'Thương mại', 'PLS-SEM cho mô hình khám phá', 'KHÁC SEM truyền thống (CB-SEM)'],
        ['JASP/jamovi', 'Open-source', 'GUI, miễn phí, dễ học', 'Tính năng SEM còn hạn chế'],
    ]
)

# 8. Năm điểm cần lưu ý khi báo cáo SEM
H('8. Năm điểm cần lưu ý khi báo cáo SEM', level=2, color=NAVY)
para(
    'Theo APA Style Guide và Boomsma (2000), một báo cáo SEM '
    'chuẩn phải có năm thành phần.', indent=False
)
para(
    'ĐIỂM 1 — Mô hình lý thuyết. Vẽ rõ đường dẫn dự kiến và lý '
    'do mỗi đường dựa trên cơ sở lý thuyết. KHÔNG fit data trước '
    'rồi mới tìm lý thuyết — vi phạm HARKing (Hypothesizing '
    'After Results Known).'
)
para(
    'ĐIỂM 2 — Cỡ mẫu và phương pháp ước lượng. Báo cáo n, '
    'phương pháp (Maximum Likelihood là chuẩn), xử lý dữ liệu '
    'thiếu (FIML hoặc multiple imputation), và biến đổi nếu '
    'có.'
)
para(
    'ĐIỂM 3 — Hệ số. Báo cáo CẢ β chuẩn hóa và β chưa chuẩn '
    'hóa với standard error (SE) và p-value. KTC 95% cho hệ '
    'số trung gian (mediation) qua bootstrap.'
)
para(
    'ĐIỂM 4 — Fit indices. Báo cáo TỐI THIỂU 4 chỉ số: '
    'χ²(df, p), CFI, RMSEA (với KTC 90%), SRMR. So sánh với '
    'ngưỡng Hu & Bentler (1999).'
)
para(
    'ĐIỂM 5 — R² cho biến phụ thuộc. Đo lường explained '
    'variance. Phân loại theo Cohen (1988): R² ≥ 0,02 nhỏ; '
    '≥ 0,13 trung bình; ≥ 0,26 lớn.'
)

# 9. CÂU TRẢ LỜI
H('9. CÂU TRẢ LỜI tóm gọn', level=2, color=NAVY)
blue_run('SEM là gì?', bold=True)
blue_run(
    'Mô hình phương trình cấu trúc — kết hợp phân tích nhân '
    'tố khẳng định (CFA) cho biến tiềm ẩn + hồi quy đường '
    'dẫn (path analysis) cho mối quan hệ. Cho phép kiểm định '
    'đồng thời nhiều mối quan hệ + biến trung gian + sai số '
    'đo trong một mô hình tổng thể.'
)
blue_run('β chuẩn hóa (standardized) vs β chưa chuẩn hóa:', bold=True)
blue_run(
    '• β chưa chuẩn hóa (B): X tăng 1 đơn vị gốc → Y tăng B '
    'đơn vị gốc.\n'
    '• β chuẩn hóa (β): X tăng 1 SD → Y tăng β SD. Không có '
    'đơn vị → SO SÁNH ĐƯỢC giữa các biến đo bằng đơn vị '
    'khác nhau.\n'
    '• Quy đổi: β = B × (SD_X / SD_Y).\n'
    '• ƯU TIÊN báo cáo β chuẩn hóa khi muốn so sánh cường '
    'độ tác động.'
)
blue_run('Fit indices chuẩn vàng (Hu & Bentler, 1999):', bold=True)
blue_run(
    '• CFI ≥ 0,95 (good); ≥ 0,90 (acceptable)\n'
    '• TLI ≥ 0,95 (good); ≥ 0,90 (acceptable)\n'
    '• RMSEA ≤ 0,06 (good); ≤ 0,08 (acceptable). KTC 90% '
    'phải chứa 0,05.\n'
    '• SRMR ≤ 0,08 (good); ≤ 0,10 (acceptable)\n'
    '• χ²/df < 3 (good); < 5 (acceptable). χ² đơn lẻ '
    'KHÔNG đáng tin với cỡ mẫu lớn.'
)
blue_run('Phối hợp 2 chỉ số (Hu & Bentler):', bold=True)
blue_run(
    'CFI ≥ 0,95 KẾT HỢP SRMR ≤ 0,08, hoặc CFI ≥ 0,95 KẾT '
    'HỢP RMSEA ≤ 0,06 — giảm Type I + Type II errors.'
)
blue_run('Cỡ mẫu khuyến nghị:', bold=True)
blue_run(
    'Tối thiểu n = 200 cho mô hình đơn giản; n ≥ 400 cho mô '
    'hình phức tạp; tỷ lệ N:q ratio ≥ 5:1 (Bentler & Chou, '
    '1987). Chương 3 luận án có n = 1.352 — RẤT TỐT cho '
    'mô hình tích hợp 6 biến tiềm ẩn.'
)
blue_run('Áp dụng cho chương 3 luận án:', bold=True)
blue_run(
    '• ALHT → RLLATQ: β = 0,510 (mạnh)\n'
    '• ALHT → RLLAXH: β = 0,490 (mạnh)\n'
    '• TTr → RLLATQ: β = −0,455 (NGANG ALHT ~89%)\n'
    '• TTr → RLLAXH: β = −0,415 (NGANG ALHT ~85%)\n'
    '• BPĐP → RLLATQ: β = 0,749 (CỰC LỚN — nhưng fit '
    'CHƯA TỐT, cần phân tách)\n'
    '• YTNC riêng: β = 0,747; YTBV riêng: β = 0,352\n'
    '• Tích hợp YTNC + YTBV: R² = 0,598 (effect size LỚN '
    'theo Cohen 1988)'
)

# 10. TLTK
H('10. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Bentler, P. M., & Chou, C. P. (1987). Practical issues in structural modeling. Sociological Methods & Research, 16(1), 78–117. https://doi.org/10.1177/0049124187016001004',
    'Boomsma, A. (2000). Reporting analyses of covariance structures. Structural Equation Modeling, 7(3), 461–483. https://doi.org/10.1207/S15328007SEM0703_6',
    'Carver, C. S. (1997). You want to measure coping but your protocol\'s too long: Consider the Brief COPE. International Journal of Behavioral Medicine, 4(1), 92–100. https://doi.org/10.1207/s15327558ijbm0401_6',
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
    'Kline, R. B. (2015). Principles and practice of structural equation modeling (4th ed.). The Guilford Press.',
    'Zheng, G. F., & Peng, H. Y. (2025). The effects of social media addiction, academic stress, and sleep quality on anxiety symptoms: A cross-sectional study of Chinese vocational students. Psychology Research and Behavior Management, 18, 1571–1584. https://doi.org/10.2147/PRBM.S522652',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
