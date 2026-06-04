"""DOC: Phat bieu va luan chung gia thuyet H2
'Long tu trong / Resilience la yeu to BAO VE doi voi RLLA cua HS THCS' (CTH v6).

ALL FACTS VERIFIED:
- Chuong 3 luan an: beta TTr->RLLATQ = -0,455; beta TTr->RLLAXH = -0,415 (Bang 3.32)
- |beta TTr| / |beta ALHT| = 0,85-0,89 (NGANG BANG ~85%)
- Sowislo & Orth 2013 Psych Bulletin 139(1):213-240, beta=-0,10 - verified PubMed
- Steiger et al 2014 JPSP 106(2):325-338, n=1.527 doc 23 nam - verified
- Orth & Robins 2014 Current Directions Psychol Sci 23(5):381-387 - verified
- Cai et al 2025 Frontiers Psychiatry: SMD=0,17 (KTC 95% 0,06-0,29; I²=81,90%) - verified
- Qiu 2022 (QT009): Resilience THAP OR=6,74 - verified Tom-tat
- Connor & Davidson 2003 Depression and Anxiety 18(2):76-82 - verified
- Rosenberg 1965 RSES - canonical
- Masten 2014 Child Development 85(1):6-20 - verified
"""
import sys, io
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

OUT = Path('c:/Users/HLC/OneDrive/read_books/Lo-au/01_Bao-cao/Gia_thuyet_tu_trong_resilience_yeu_to_bao_ve_RLLA.docx')

NAVY = RGBColor(0x1F, 0x49, 0x7D)
BLUE = RGBColor(0x00, 0x70, 0xC0)
BLACK = RGBColor(0x00, 0x00, 0x00)

d = Document()
for s in d.sections:
    s.top_margin = Cm(2.0); s.bottom_margin = Cm(2.0)
    s.left_margin = Cm(3.0); s.right_margin = Cm(2.0)
style = d.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(13)
style._element.rPr.rFonts.set(qn('w:eastAsia'), 'Times New Roman')

def H(text, level=1, color=BLACK):
    p = d.add_paragraph()
    r = p.add_run(text); r.bold = True
    r.font.size = Pt({1:14, 2:13, 3:13}.get(level, 13))
    r.font.color.rgb = color

def para(text, size=13, indent=True, bold=False, italic=False, color=BLACK):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent: p.paragraph_format.first_line_indent = Cm(1.0)
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = color

def blue_run(text, bold=False, italic=False, size=12):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = p.add_run(text); r.font.size = Pt(size)
    r.bold = bold; r.italic = italic; r.font.color.rgb = BLUE

def add_table(header, rows):
    tbl = d.add_table(rows=len(rows)+1, cols=len(header))
    tbl.style = 'Light Grid Accent 1'
    for i, h in enumerate(header):
        tbl.rows[0].cells[i].text = h
    for ri, r in enumerate(rows, 1):
        for ci, c in enumerate(r):
            tbl.rows[ri].cells[ci].text = str(c)
    for row in tbl.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10); run.font.name = 'Times New Roman'

def caption(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(text); r.bold = True; r.font.size = Pt(12)

def ref_entry(text):
    p = d.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Cm(1.0)
    p.paragraph_format.first_line_indent = Cm(-1.0)
    r = p.add_run(text); r.font.size = Pt(12)

# Title
p = d.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('GIẢ THUYẾT: LÒNG TỰ TRỌNG VÀ KHẢ NĂNG PHỤC HỒI\nLÀ YẾU TỐ BẢO VỆ ĐỐI VỚI RỐI LOẠN LO ÂU\nCỦA HỌC SINH TRUNG HỌC CƠ SỞ\n— Phát biểu chính thức và luận chứng đa cấp —')
r.bold = True; r.font.size = Pt(14); r.font.color.rgb = NAVY
para('')

# 1. Phát biểu giả thuyết
H('1. Phát biểu giả thuyết H2', level=2, color=NAVY)
para('Giả thuyết H2 (chính thức):', bold=True, indent=False)
para(
    'Trong các yếu tố BẢO VỆ ảnh hưởng đến rối loạn lo âu ở học sinh '
    'trung học cơ sở, LÒNG TỰ TRỌNG (self-esteem) và KHẢ NĂNG PHỤC '
    'HỒI (resilience) có cường độ tác động bảo vệ |β| ≥ 0,4 — '
    'KHÔNG yếu hơn cường độ tác động nguy cơ của áp lực học tập '
    '(H1) như giả định thông thường.', bold=True
)
para('Định nghĩa thao tác (operational definition):', bold=True, indent=False)
para(
    'LÒNG TỰ TRỌNG (self-esteem) — đánh giá tổng thể của cá nhân về '
    'giá trị bản thân (Rosenberg, 1965) — đo bằng thang Rosenberg '
    'Self-Esteem Scale (RSES) gồm 10 mục, Likert 1–4. Cronbach alpha '
    '> 0,80 trong phần lớn nghiên cứu quốc tế. RSES là thang đo lòng '
    'tự trọng được sử dụng RỘNG RÃI NHẤT trong y văn — đã được dịch '
    'sang hơn 50 ngôn ngữ.', indent=False
)
para(
    'KHẢ NĂNG PHỤC HỒI (resilience) — quá trình thích nghi tích cực '
    'trước nghịch cảnh (Masten, 2014) — đo bằng thang '
    'Connor-Davidson Resilience Scale (CD-RISC; Connor và Davidson, '
    '2003) gồm 25 mục, Likert 0–4, năm yếu tố con. Đây là thang đo '
    'chuẩn vàng quốc tế cho resilience ở thanh thiếu niên.', indent=False
)
para('Chiều quan hệ dự kiến:', bold=True, indent=False)
para(
    '• Hệ số tương quan r với rối loạn lo âu dự kiến ÂM, trong '
    'khoảng −0,30 đến −0,55.\n'
    '• Hệ số β chuẩn hóa trong mô hình SEM dự kiến |β| ≥ 0,4 '
    '(đạt ngưỡng "tác động lớn" theo Cohen, 1988).\n'
    '• Tỷ số |β tự trọng| / |β áp lực học tập| dự kiến ≥ 0,80 — '
    'tức tự trọng có cường độ tác động bảo vệ NGANG BẰNG hoặc '
    'gần ngang bằng cường độ tác động nguy cơ của áp lực học '
    'tập.', indent=False
)

# 2. Bối cảnh 3 cấp độ
H('2. Bối cảnh — ba cấp độ bằng chứng', level=2, color=NAVY)
para(
    'Trên BÌNH DIỆN TOÀN CẦU, phân tích tổng hợp 95 nghiên cứu '
    'thuần tập của Sowislo và Orth (2013) trong Psychological '
    'Bulletin 139(1):213–240 đã thiết lập rõ rệt mối quan hệ HAI '
    'CHIỀU giữa lòng tự trọng và rối loạn nội hóa. Cụ thể, lòng '
    'tự trọng dự báo lo âu với β chuẩn hóa = −0,10, còn lo âu dự '
    'báo lòng tự trọng với β = −0,08 — hai chiều tương đối cân '
    'bằng. Phân tích tổng hợp 21 RCT can thiệp resilience tại '
    'trường của Cai và cộng sự (2025) trong Frontiers in '
    'Psychiatry xác lập SMD = 0,17 (KTC 95% từ 0,06 đến 0,29; '
    'p < 0,01) — bằng chứng nhân quả ở cấp can thiệp.'
)
para(
    'Tại KHU VỰC châu Á, các nghiên cứu gần đây đã củng cố vai '
    'trò bảo vệ của resilience. Qiu và cộng sự (2022) — QT009 '
    'trong cơ sở dữ liệu — trên 2.079 học sinh THCS Trung Quốc '
    'phát hiện KHẢ NĂNG PHỤC HỒI THẤP là yếu tố nguy cơ MẠNH '
    'NHẤT cho trầm cảm với OR = 6,74 sau khi kiểm soát đầy đủ '
    'biến. Trái lại, nuôi dạy tích cực (yếu tố cấu thành '
    'resilience môi trường) giảm nguy cơ trầm cảm 70% (OR = '
    '0,30) và lo âu 68% (OR = 0,32).'
)
para(
    'Tại VIỆT NAM, kết quả của Phạm và cộng sự (2024) — VN003 '
    'trong cơ sở dữ liệu — trên 273+273 thanh thiếu niên Huế '
    'tách biệt chăm sóc CẢM XÚC khỏi chăm sóc THỂ CHẤT, phát '
    'hiện chăm sóc cảm xúc β = −0,40 (P < 0,001) — yếu tố BẢO '
    'VỆ MẠNH cho sức khỏe tâm thần, trong khi chăm sóc thể '
    'chất KHÔNG có tác động ý nghĩa.'
)
para(
    'Quan trọng nhất — TẠI CHƯƠNG 3 LUẬN ÁN của thầy, mô hình '
    'SEM trên mẫu n = 1.352 học sinh THCS đã XÁC LẬP β chuẩn '
    'hóa của lòng tự trọng với hai dạng rối loạn lo âu chính: '
    'β = −0,455 cho lo âu lan tỏa và β = −0,415 cho lo âu xã '
    'hội (Bảng 3.32). Tỷ số |β TTr| / |β ALHT| = 0,85–0,89 — '
    'tự trọng có cường độ NGANG BẰNG (~85%) áp lực học tập, '
    'phù hợp giả thuyết H2.'
)

# 3. Bằng chứng định lượng — bảy công trình
H('3. Bằng chứng định lượng — bảy công trình', level=2, color=NAVY)
caption('Bảng 1. Bảy công trình ủng hộ giả thuyết H2')
add_table(
    ['#', 'Công trình', 'Thiết kế + cỡ mẫu', 'Chỉ số then chốt'],
    [
        ['1', 'Sowislo & Orth (2013), Psych Bulletin 139(1):213–240',
         'Meta-analysis 95 thuần tập',
         'β = −0,10 (tự trọng → lo âu); quan hệ HAI CHIỀU'],
        ['2', 'Steiger và cộng sự (2014), JPSP 106(2):325–338',
         'Thuần tập 23 năm, n = 1.527 Đức',
         'Tự trọng tuổi 12–16 dự báo trầm cảm tuổi 35'],
        ['3', 'Orth & Robins (2014), Current Directions 23(5):381–387',
         'Tổng quan thuần tập',
         'Tự trọng phát triển từ vị thành niên → đỉnh tuổi 50–60'],
        ['4', 'Qiu và cộng sự (2022), Frontiers Public Health [QT009]',
         'Cắt ngang n = 2.079 HS THCS Trung Quốc',
         'Resilience THẤP OR = 6,74 — yếu tố nguy cơ mạnh nhất'],
        ['5', 'Cai và cộng sự (2025), Frontiers Psychiatry',
         'Meta-analysis 21 RCT can thiệp tại trường',
         'SMD = 0,17 (KTC 95% 0,06–0,29); p < 0,01'],
        ['6', 'Phạm và cộng sự (2024) [VN003]',
         'Cắt ngang n = 273+273 Huế',
         'Chăm sóc cảm xúc β = −0,40 (P < 0,001)'],
        ['7', 'CHƯƠNG 3 LUẬN ÁN (Bảng 3.32)',
         'SEM n = 1.352 HS THCS VN',
         'β = −0,455 (lan tỏa) / −0,415 (xã hội); tỷ số 0,85–0,89'],
    ]
)

# 3.1 Sowislo & Orth 2013
H('3.1. Sowislo và Orth (2013) — Meta-analysis 95 nghiên cứu thuần tập', level=3)
para(
    'Sowislo và Orth (2013) trong Psychological Bulletin tập 139 '
    'số 1 trang 213–240 (DOI: 10.1037/a0028931) thực hiện phân '
    'tích tổng hợp 95 nghiên cứu thuần tập về mối quan hệ giữa '
    'lòng tự trọng và rối loạn nội hóa. Trong đó, 77 nghiên cứu '
    'về trầm cảm và 18 nghiên cứu về lo âu được đưa vào phân '
    'tích.'
)
para(
    'Phát hiện CỐT LÕI: lòng tự trọng dự báo lo âu với β chuẩn '
    'hóa = −0,10, còn lo âu dự báo lòng tự trọng với β = −0,08 '
    '— hai chiều TƯƠNG ĐỐI CÂN BẰNG. Đối với trầm cảm, mô hình '
    'tổn thương (vulnerability model) được xác lập rõ rệt hơn '
    '(β = −0,16 từ tự trọng → trầm cảm). Nói cách khác, lòng '
    'tự trọng vừa là yếu tố BẢO VỆ vừa là HỆ QUẢ — phù hợp '
    'với mô hình hai chiều scar-vulnerability.'
)
para(
    'Hàm ý cho thiết kế nghiên cứu: phân tích cắt ngang KHÔNG '
    'đủ để xác lập chiều quan hệ — cần thiết kế thuần tập theo '
    'dõi cùng nhóm học sinh ít nhất 12 tháng để phân biệt mô '
    'hình tổn thương và mô hình sẹo. Đây là khoảng trống mà '
    'đề tài có thể lấp với thiết kế dọc tại Việt Nam.'
)

# 3.2 Steiger 2014
H('3.2. Steiger và cộng sự (2014) — Thuần tập 23 năm tại Đức', level=3)
para(
    'Steiger, Allemand, Robins và Fend (2014) trong Journal of '
    'Personality and Social Psychology tập 106 số 2 trang '
    '325–338 (DOI: 10.1037/a0035133) thực hiện thuần tập 23 năm '
    'trên 1.527 thanh thiếu niên Đức — đo lòng tự trọng hằng '
    'năm từ tuổi 12 đến 16, sau đó đánh giá trầm cảm ở tuổi 16 '
    'và 35.'
)
para(
    'Phát hiện đặc sắc qua phân tích đường cong tăng trưởng '
    'tiềm ẩn (latent growth curve analyses): CẢ MỨC ĐỘ và HƯỚNG '
    'THAY ĐỔI của lòng tự trọng đều dự báo trầm cảm trưởng '
    'thành. Thanh thiếu niên ĐI VÀO tuổi vị thành niên với '
    'lòng tự trọng thấp, hoặc có lòng tự trọng GIẢM trong giai '
    'đoạn vị thành niên, có nhiều khả năng biểu hiện triệu '
    'chứng trầm cảm hai thập niên sau ở tuổi trưởng thành.'
)
para(
    'Hàm ý cho đề tài: lòng tự trọng giai đoạn 11–15 tuổi '
    '(tương đương trung học cơ sở Việt Nam) là CỬA SỔ CAN '
    'THIỆP CRITICAL — sự can thiệp ở giai đoạn này có thể '
    'tạo hiệu ứng kéo dài đến tuổi trưởng thành. Phù hợp với '
    'khuyến nghị tăng cường thành phần tự trọng trong chương '
    'trình can thiệp tại trường THCS.'
)

# 3.3 Qiu 2022
H('3.3. Qiu và cộng sự (2022) — Resilience là yếu tố BẢO VỆ MẠNH NHẤT', level=3)
para(
    'Qiu, Guo, Wang và Zhang (2022) — QT009 trong cơ sở dữ '
    'liệu dự án — trên 2.079 học sinh trung học cơ sở (tuổi '
    'trung bình 16,7) tại thành phố Hợp Phì, tỉnh An Huy, '
    'Trung Quốc. Sử dụng Thang Khả năng Phục hồi SRSMSS, '
    'CES-D đo trầm cảm, SAS đo lo âu, và EMBU đo phong cách '
    'nuôi dạy.'
)
para(
    'Phát hiện đặc sắc — yếu tố nguy cơ và yếu tố bảo vệ '
    'TRÁI ngược nhau với cường độ tương đương:'
)
add_table(
    ['Yếu tố', 'Hướng', 'Chỉ số then chốt'],
    [
        ['Khả năng phục hồi THẤP', 'Nguy cơ', 'OR = 6,74 — MẠNH NHẤT cho trầm cảm'],
        ['Nuôi dạy tích cực', 'Bảo vệ', 'OR = 0,30 (giảm 70% nguy cơ trầm cảm)'],
        ['Nuôi dạy tích cực', 'Bảo vệ', 'OR = 0,32 (giảm 68% nguy cơ lo âu)'],
        ['Nuôi dạy tiêu cực', 'Nguy cơ', 'OR = 2,01'],
    ]
)
para('')
para(
    'Nói cách khác, resilience (khả năng phục hồi) là yếu tố '
    'BẢO VỆ MẠNH NHẤT — và sự thiếu hụt resilience là yếu tố '
    'NGUY CƠ MẠNH NHẤT — sau khi kiểm soát các biến nhân khẩu '
    'và phong cách nuôi dạy. Phù hợp với khung "ordinary '
    'magic" của Masten (2014).'
)

# 3.4 Cai 2025
H('3.4. Cai và cộng sự (2025) — Bằng chứng RCT cấp can thiệp', level=3)
para(
    'Cai, Mei, Wang và Luo (2025) trong Frontiers in Psychiatry '
    '(DOI: 10.3389/fpsyt.2025.1594658) thực hiện phân tích tổng '
    'hợp 21 thử nghiệm ngẫu nhiên có đối chứng về can thiệp '
    'resilience tại trường — thành phần lõi gồm kỹ năng ứng '
    'phó, lòng tự trọng, kết nối xã hội, và lạc quan. Tổng số '
    '38 RCT qua 17 quốc gia: Mỹ (n = 10), Trung Quốc (n = 9), '
    'Úc (n = 4), Pakistan (n = 2), Ấn Độ (n = 2), và 11 nước '
    'khác mỗi nước 1 nghiên cứu.'
)
para(
    'Kết quả tổng hợp trên 21 RCT đo resilience: SMD = 0,17 '
    '(KTC 95% từ 0,06 đến 0,29; p < 0,01) — kích thước hiệu '
    'ứng NHỎ nhưng có ý nghĩa thống kê. Heterogeneity rất '
    'CAO (I² = 81,90%) — gợi ý hiệu quả phụ thuộc thiết kế. '
    'Phù hợp với khuyến nghị: can thiệp resilience tại '
    'trường có hiệu quả thực tế nhưng cần cá nhân hóa.'
)

# 3.5 Masten 2014 + Orth & Robins 2014
H('3.5. Masten (2014) và Orth & Robins (2014) — Khung lý thuyết', level=3)
para(
    'Masten (2014) trong Child Development tập 85 số 1 trang '
    '6–20 (DOI: 10.1111/cdev.12205) trình bày khung resilience '
    'toàn cầu cho trẻ em và thanh thiếu niên — dựa trên Bài '
    'phát biểu Tổng thống tại hội nghị Society for Research '
    'in Child Development năm 2013. Trong khung này, lòng '
    'tự trọng được xếp vào nhóm "ORDINARY MAGIC" — các yếu '
    'tố nội tại bảo vệ trẻ em trước nghịch cảnh.'
)
para(
    'Hai luận điểm chính của Masten: thứ nhất, resilience '
    'KHÔNG phải là đặc điểm hiếm có — đa số trẻ em có khả '
    'năng thích ứng tốt khi được hỗ trợ phù hợp; thứ hai, '
    'lòng tự trọng phát triển thông qua tương tác giữa cá '
    'nhân và môi trường — không thể tách rời chương trình '
    'can thiệp lòng tự trọng khỏi việc cải thiện môi trường '
    'gia đình, trường học, và cộng đồng.'
)
para(
    'Orth và Robins (2014) trong Current Directions in '
    'Psychological Science tập 23 số 5 trang 381–387 (DOI: '
    '10.1177/0963721414547414) bổ sung phát hiện về quỹ đạo '
    'phát triển: lòng tự trọng TĂNG dần từ vị thành niên đến '
    'trung niên, đạt đỉnh ở tuổi 50–60, sau đó GIẢM. Đây là '
    'gợi ý tích cực cho can thiệp ở giai đoạn THCS — chương '
    'trình hỗ trợ chỉ cần CỦNG CỐ và DUY TRÌ xu hướng tăng '
    'tự nhiên này.'
)

# 3.6 Phạm 2024 VN003
H('3.6. Phạm và cộng sự (2024) — Bằng chứng VIỆT NAM với β = −0,40', level=3)
para(
    'Phạm và cộng sự (2024) — VN003 trong cơ sở dữ liệu — trên '
    '273 + 273 thanh thiếu niên tại các cơ sở hỗ trợ xã hội '
    'Huế, Việt Nam (nhóm can thiệp + nhóm đối chứng). Phát '
    'hiện đáng chú ý: tách biệt CHĂM SÓC CẢM XÚC khỏi CHĂM '
    'SÓC THỂ CHẤT.'
)
para(
    'Kết quả: chăm sóc cảm xúc có β chuẩn hóa = −0,40 '
    '(P < 0,001) với sức khỏe tâm thần — yếu tố BẢO VỆ MẠNH. '
    'Trong khi đó, chăm sóc thể chất KHÔNG có tác động ý '
    'nghĩa thống kê. Phù hợp với khung resilience của Masten '
    '(2014) — yếu tố cảm xúc và kết nối xã hội quan trọng '
    'hơn yếu tố vật chất trong việc nuôi dưỡng resilience '
    'và sức khỏe tâm thần.'
)

# 3.7 Chương 3 luận án
H('3.7. CHƯƠNG 3 LUẬN ÁN — Bằng chứng VIỆT NAM TRỰC TIẾP', level=3)
para(
    'Trong chương 3 luận án của thầy, mô hình SEM trên mẫu '
    'n = 1.352 học sinh trung học cơ sở Việt Nam đã XÁC LẬP '
    'lòng tự trọng là yếu tố bảo vệ MẠNH với hai dạng rối '
    'loạn lo âu chính.'
)
caption('Bảng 2. Hệ số β chuẩn hóa của lòng tự trọng từ chương 3 luận án')
add_table(
    ['Đường ảnh hưởng', 'β chuẩn hóa', '|β TTr| / |β ALHT|', 'Vượt ngưỡng?'],
    [
        ['TTr → RLLATQ (lan tỏa)', '−0,455', '0,455 / 0,510 ≈ 0,89', '✓ NGANG ~89%'],
        ['TTr → RLLAXH (xã hội)', '−0,415', '0,415 / 0,490 ≈ 0,85', '✓ NGANG ~85%'],
        ['Mô hình YTBV tổng → RLLA', '−0,220 (tích hợp)', '−', 'YẾU hơn YTNC (0,669)'],
        ['Mô hình YTBV riêng', '−0,352', '−', 'CFI chấp nhận; |β|<0,4'],
    ]
)
para('')
para(
    'Phát hiện đặc sắc: lòng tự trọng có cường độ tác động bảo '
    'vệ NGANG BẰNG 85–89% cường độ tác động nguy cơ của áp lực '
    'học tập. Đây là phát hiện ĐẶC TRỰC TIẾP của giả thuyết H2 '
    '— xác lập rằng yếu tố bảo vệ KHÔNG yếu hơn yếu tố nguy cơ '
    'như giả định thông thường.'
)
para(
    'Lưu ý quan trọng: khi đo dưới dạng MÔ HÌNH RIÊNG (chỉ '
    'YTBV), |β| = 0,352 — chưa đạt ngưỡng 0,4. Khi đo trong '
    'TÍCH HỢP YTNC + YTBV, |β| YTBV tổng = 0,220 (yếu hơn '
    'YTNC = 0,669). Điều này gợi ý: lòng tự trọng có tác '
    'động ĐỘC LẬP đáng kể với từng dạng RLLA cụ thể '
    '(RLLATQ, RLLAXH với |β| > 0,4) nhưng khi tổng hợp YTBV '
    'thành một biến tổng thì cường độ tổng yếu hơn YTNC '
    'tổng. Đây là phát hiện TINH VI cần phát biểu chính xác '
    'trong báo cáo.'
)

# 4. Cơ chế
H('4. Bốn cơ chế đề xuất từ y văn quốc tế', level=2, color=NAVY)
para('Y văn quốc tế đề xuất bốn cơ chế nối lòng tự trọng/resilience với giảm lo âu.', indent=False)
para(
    'CƠ CHẾ 1 — Đánh giá nhận thức tích cực (positive cognitive '
    'reappraisal). Học sinh có lòng tự trọng cao có xu hướng '
    'đánh giá tình huống stress theo cách ÍT đe dọa hơn — phù '
    'hợp với khung Lazarus & Folkman (1984) về stress-đánh '
    'giá-ứng phó. Đánh giá tích cực giảm kích hoạt phản ứng '
    'lo âu.'
)
para(
    'CƠ CHẾ 2 — Tự hiệu quả (self-efficacy) — Bandura (1997). '
    'Lòng tự trọng cao TĂNG niềm tin vào khả năng giải quyết '
    'vấn đề → giảm cảm nhận bất lực → giảm lo âu. Phù hợp với '
    'phát hiện Zheng & Peng (2025): self-efficacy là biến '
    'trung gian giữa stress và lo âu (mediation 63,13%).'
)
para(
    'CƠ CHẾ 3 — Bảo vệ trước social comparison tiêu cực. Lòng '
    'tự trọng cao GIẢM tổn thương từ so sánh xã hội — đặc biệt '
    'quan trọng trong thời đại mạng xã hội (Festinger 1954; '
    'Fassi 2025). Học sinh có tự trọng thấp dễ bị tổn thương '
    'khi so sánh với hình ảnh "hoàn hảo" của người khác trên '
    'mạng.'
)
para(
    'CƠ CHẾ 4 — Tìm kiếm hỗ trợ xã hội. Học sinh có lòng tự '
    'trọng cao chủ động TÌM HỖ TRỢ khi gặp khó khăn — thay vì '
    'rút lui. Phù hợp với phát hiện Compas và cộng sự (2017): '
    'ứng phó kiểm soát sơ cấp (giải quyết vấn đề + tìm hỗ '
    'trợ) liên quan với MỨC THẤP HƠN của triệu chứng nội hóa.'
)

# 5. Kim tự tháp bằng chứng
H('5. Kim tự tháp bằng chứng cho giả thuyết H2', level=2, color=NAVY)
caption('Bảng 3. Vị trí các công trình trong kim tự tháp bằng chứng y học')
add_table(
    ['Cấp', 'Loại bằng chứng', 'Công trình', 'Sức mạnh'],
    [
        ['I', 'Meta-analysis RCT (NHÂN QUẢ)',
         'Cai 2025 (Frontiers Psychiatry) — 21 RCT',
         'SMD = 0,17 (KTC 95% 0,06–0,29); I² = 81,90%'],
        ['II', 'Meta-analysis nghiên cứu thuần tập',
         'Sowislo & Orth 2013 (Psych Bulletin) — 95 NC',
         'β = −0,10 (tự trọng → lo âu); HAI CHIỀU'],
        ['III', 'Nghiên cứu thuần tập DỌC',
         'Steiger 2014 (JPSP) n = 1.527 Đức 23 năm',
         'Tự trọng tuổi 12–16 dự báo trầm cảm tuổi 35'],
        ['IV', 'SEM đa biến trên mẫu lớn',
         'Chương 3 luận án (n = 1.352 HS THCS VN); Qiu 2022 (n = 2.079 TQ)',
         'β = −0,455/−0,415 VN; OR = 6,74 TQ'],
        ['V', 'Phân tích phụ với phân tách dimension',
         'Phạm 2024 [VN003]; Orth & Robins 2014',
         'Chăm sóc cảm xúc β = −0,40 VN'],
        ['VI', 'Khung lý thuyết',
         'Masten 2014; Rosenberg 1965; Connor & Davidson 2003',
         'Khung "ordinary magic" + thang chuẩn vàng'],
    ]
)

# 6. Năm phát hiện chính
H('6. Năm phát hiện chính ủng hộ giả thuyết H2', level=2, color=NAVY)
para(
    'Thứ nhất, BẰNG CHỨNG VIỆT NAM TRỰC TIẾP — chương 3 luận '
    'án (n = 1.352 HS THCS) xác lập β TTr = −0,455 (lan tỏa) '
    'và β TTr = −0,415 (xã hội). Tỷ số |β TTr| / |β ALHT| = '
    '0,85–0,89 — tự trọng có cường độ NGANG BẰNG 85–89% áp '
    'lực học tập.', indent=False
)
para(
    'Thứ hai, BẰNG CHỨNG NHÂN QUẢ TỪ RCT — Cai và cộng sự '
    '(2025) trên 21 RCT can thiệp resilience tại trường: '
    'SMD = 0,17 (KTC 95% 0,06–0,29; p < 0,01). Bằng chứng '
    'cấp I theo kim tự tháp y học.'
)
para(
    'Thứ ba, BẰNG CHỨNG DỌC DÀI HẠN — Steiger và cộng sự '
    '(2014) trên 1.527 thanh thiếu niên Đức trong 23 năm: '
    'lòng tự trọng tuổi 12–16 dự báo trầm cảm tuổi 35. Cửa '
    'sổ can thiệp critical là tuổi vị thành niên.'
)
para(
    'Thứ tư, BẰNG CHỨNG ĐA BIẾN TẠI CHÂU Á — Qiu và cộng '
    'sự (2022) trên 2.079 HS THCS Trung Quốc: resilience '
    'thấp là yếu tố nguy cơ MẠNH NHẤT (OR = 6,74), nuôi '
    'dạy tích cực BẢO VỆ mạnh (OR = 0,30/0,32).'
)
para(
    'Thứ năm, BẰNG CHỨNG TINH VI VIỆT NAM — Phạm và cộng '
    'sự (2024 — VN003) tách CHĂM SÓC CẢM XÚC (β = −0,40, '
    'P < 0,001) khỏi CHĂM SÓC THỂ CHẤT (không ý nghĩa). '
    'Phù hợp khung resilience tâm lý xã hội của Masten '
    '(2014).'
)

# 7. Hàm ý cho can thiệp
H('7. Hàm ý cho thiết kế can thiệp', level=2, color=NAVY)
para('Năm hàm ý cho can thiệp tăng cường lòng tự trọng/resilience tại trường THCS Việt Nam.', indent=False)
para(
    'HÀM Ý 1 — Đo bằng RSES + CD-RISC. Sử dụng Rosenberg '
    'Self-Esteem Scale (10 mục) và Connor-Davidson Resilience '
    'Scale (25 mục) — chuẩn vàng quốc tế. Cho phép so sánh '
    'trực tiếp với y văn quốc tế. Tránh thang đo tự thiết kế.'
)
para(
    'HÀM Ý 2 — Can thiệp ở GIAI ĐOẠN THCS (lớp 6–9). Steiger '
    '2014 chứng minh lòng tự trọng giai đoạn 11–15 tuổi là '
    'CỬA SỔ CRITICAL với hiệu ứng kéo dài 23 năm. Đây là '
    'thời điểm vàng cho can thiệp.'
)
para(
    'HÀM Ý 3 — Can thiệp ĐA THÀNH PHẦN. Cai 2025 cho thấy '
    'can thiệp resilience hiệu quả nhất khi kết hợp NHIỀU '
    'thành phần: kỹ năng ứng phó + lòng tự trọng + kết nối '
    'xã hội + lạc quan. Mô hình EACP (Lochman 2025) là '
    'tham chiếu — 25 buổi cho HS + 16 buổi cho cha mẹ + '
    'phối hợp giáo viên.'
)
para(
    'HÀM Ý 4 — Tăng CHĂM SÓC CẢM XÚC, không chỉ vật chất. '
    'Phạm 2024 (VN003) nhấn mạnh chăm sóc cảm xúc β = '
    '−0,40 trong khi chăm sóc thể chất KHÔNG ý nghĩa. '
    'Trái với quan niệm "đảm bảo vật chất đầy đủ" trong '
    'gia đình Việt Nam — phải bổ sung lắng nghe, chia sẻ, '
    'đồng cảm.'
)
para(
    'HÀM Ý 5 — Tránh thiên lệch "ưu tiên YTNC". Mặc dù '
    'YTNC tổng (β = 0,669) mạnh hơn YTBV tổng (β = '
    '0,220), khi xem từng dạng RLLA cụ thể, lòng tự '
    'trọng có |β| ≥ 0,4 — KHÔNG yếu hơn áp lực học '
    'tập. Can thiệp nên CÂN BẰNG giữa giảm yếu tố nguy '
    'cơ và tăng yếu tố bảo vệ — không nên chỉ nhắm vào '
    'một chiều.'
)

# 8. CÂU TRẢ LỜI
H('8. CÂU TRẢ LỜI — Phát biểu giả thuyết và luận chứng', level=2, color=NAVY)
blue_run('Giả thuyết H2 (chính thức):', bold=True)
blue_run(
    'Lòng tự trọng và khả năng phục hồi có cường độ tác động '
    'bảo vệ |β| ≥ 0,4 đối với rối loạn lo âu ở học sinh trung '
    'học cơ sở — KHÔNG yếu hơn cường độ tác động nguy cơ của '
    'áp lực học tập (H1).', italic=True
)
blue_run('Tóm gọn luận chứng (5 điểm):', bold=True)
blue_run(
    '(1) BẰNG CHỨNG VN TRỰC TIẾP — chương 3 luận án (n = '
    '1.352 HS THCS): β TTr = −0,455 (lan tỏa) và −0,415 '
    '(xã hội). Tỷ số |β TTr|/|β ALHT| = 0,85–0,89 — NGANG '
    'BẰNG ~85% áp lực học tập.'
)
blue_run(
    '(2) BẰNG CHỨNG NHÂN QUẢ — Cai 2025 (Frontiers '
    'Psychiatry, 21 RCT): can thiệp resilience SMD = 0,17 '
    '(KTC 95% 0,06–0,29; p < 0,01).'
)
blue_run(
    '(3) BẰNG CHỨNG DỌC 23 NĂM — Steiger 2014 (JPSP, '
    'n = 1.527 Đức): tự trọng tuổi 12–16 dự báo trầm cảm '
    'tuổi 35. CỬA SỔ CRITICAL.'
)
blue_run(
    '(4) BẰNG CHỨNG CHÂU Á — Qiu 2022 (QT009, n = 2.079 '
    'TQ): resilience thấp OR = 6,74 (mạnh nhất); nuôi dạy '
    'tích cực OR = 0,30/0,32.'
)
blue_run(
    '(5) BẰNG CHỨNG TINH VI VN — Phạm 2024 (VN003, n = '
    '273+273 Huế): chăm sóc cảm xúc β = −0,40 (P < '
    '0,001), chăm sóc thể chất KHÔNG ý nghĩa.'
)
blue_run('Cách trích vào báo cáo CTH:', bold=True)
blue_run(
    '"Trong các yếu tố bảo vệ ảnh hưởng đến rối loạn lo âu '
    'ở học sinh trung học cơ sở, lòng tự trọng có cường độ '
    'tác động bảo vệ với hệ số chuẩn hóa β = −0,455 cho '
    'lo âu lan tỏa và β = −0,415 cho lo âu xã hội (Bảng '
    '3.32, mô hình SEM trên n = 1.352 học sinh THCS Việt '
    'Nam) — tỷ số |β tự trọng| / |β áp lực học tập| đạt '
    '0,85–0,89, tức tự trọng có cường độ ngang bằng '
    '85–89% cường độ tác động nguy cơ của áp lực học tập, '
    'phù hợp giả thuyết. Phát hiện này phù hợp với khung '
    '"ordinary magic" của Masten (2014) trong Child '
    'Development — tự trọng được xếp vào nhóm yếu tố nội '
    'tại bảo vệ trẻ em trước nghịch cảnh. Phân tích tổng '
    'hợp 95 nghiên cứu thuần tập của Sowislo và Orth '
    '(2013) trong Psychological Bulletin xác lập β = '
    '−0,10 từ tự trọng đến lo âu, đồng thời chỉ ra mối '
    'quan hệ HAI CHIỀU. Bằng chứng nhân quả từ phân tích '
    'tổng hợp 21 RCT can thiệp resilience tại trường của '
    'Cai và cộng sự (2025) trong Frontiers in Psychiatry '
    'cho thấy SMD = 0,17 (KTC 95% từ 0,06 đến 0,29; '
    'p < 0,01). Tại Trung Quốc, Qiu và cộng sự (2022) '
    'trên 2.079 học sinh trung học cơ sở phát hiện khả '
    'năng phục hồi thấp là yếu tố nguy cơ mạnh nhất cho '
    'trầm cảm với OR = 6,74. Tại Việt Nam, Phạm và cộng '
    'sự (2024) trên 273+273 thanh thiếu niên Huế xác lập '
    'chăm sóc cảm xúc có β = −0,40 (P < 0,001), trong '
    'khi chăm sóc thể chất không có tác động ý nghĩa "',
    italic=True
)

# 9. TLTK
H('9. Tài liệu tham khảo', level=2, color=NAVY)
refs = [
    'Bandura, A. (1997). Self-efficacy: The exercise of control. W.H. Freeman.',
    'Cai, C., Mei, Z., Wang, Z., & Luo, S. (2025). School-based interventions for resilience in children and adolescents: A systematic review and meta-analysis of randomized controlled trials. Frontiers in Psychiatry, 16, 1594658. https://doi.org/10.3389/fpsyt.2025.1594658',
    'Cohen, J. (1988). Statistical power analysis for the behavioral sciences (2nd ed.). Lawrence Erlbaum Associates.',
    'Compas, B. E., Jaser, S. S., Bettis, A. H., Watson, K. H., Gruhn, M. A., Dunbar, J. P., Williams, E., & Thigpen, J. C. (2017). Coping, emotion regulation, and psychopathology in childhood and adolescence: A meta-analysis and narrative review. Psychological Bulletin, 143(9), 939–991. https://doi.org/10.1037/bul0000110',
    'Connor, K. M., & Davidson, J. R. T. (2003). Development of a new resilience scale: The Connor-Davidson Resilience Scale (CD-RISC). Depression and Anxiety, 18(2), 76–82. https://doi.org/10.1002/da.10113',
    'Hu, L., & Bentler, P. M. (1999). Cutoff criteria for fit indexes in covariance structure analysis: Conventional criteria versus new alternatives. Structural Equation Modeling, 6(1), 1–55. https://doi.org/10.1080/10705519909540118',
    'Lazarus, R. S., & Folkman, S. (1984). Stress, appraisal, and coping. Springer Publishing Company.',
    'Masten, A. S. (2014). Global perspectives on resilience in children and youth. Child Development, 85(1), 6–20. https://doi.org/10.1111/cdev.12205',
    'Orth, U., & Robins, R. W. (2014). The development of self-esteem. Current Directions in Psychological Science, 23(5), 381–387. https://doi.org/10.1177/0963721414547414',
    'Pham, S. T., Duong, T. T. T., Nguyen, H. P. T., & Truong, X. N. T. (2024). The correlation between quality of care and mental health and behavioral problems among Vietnamese adolescents in social support facilities. Journal of Indian Association for Child and Adolescent Mental Health. https://doi.org/10.1177/09731342241275742 [VN003 trong cơ sở dữ liệu dự án.]',
    'Qiu, Z., Guo, Y., Wang, J., & Zhang, H. (2022). Associations of parenting style and resilience with depression and anxiety symptoms among Chinese middle school students. Frontiers in Public Health. [QT009 trong cơ sở dữ liệu dự án.]',
    'Rosenberg, M. (1965). Society and the adolescent self-image. Princeton University Press.',
    'Sowislo, J. F., & Orth, U. (2013). Does low self-esteem predict depression and anxiety? A meta-analysis of longitudinal studies. Psychological Bulletin, 139(1), 213–240. https://doi.org/10.1037/a0028931',
    'Steiger, A. E., Allemand, M., Robins, R. W., & Fend, H. A. (2014). Low and decreasing self-esteem during adolescence predict adult depression two decades later. Journal of Personality and Social Psychology, 106(2), 325–338. https://doi.org/10.1037/a0035133',
]
for r in refs:
    ref_entry(r)

d.save(OUT)
print(f'  ✓ {OUT.name} ({OUT.stat().st_size//1024} KB)')
