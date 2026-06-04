# -*- coding: utf-8 -*-
"""Tạo doc Giải trình phản hồi phản biện cho Bài 03 (mã 2026.05.14).
Định dạng (cập nhật theo yêu cầu): mỗi mục có
- "Góp ý của phản biện:" (in nghiêng đậm) + nội dung góp ý
- "Phản hồi:" (in đậm) + nội dung phản hồi (tô xanh)
- Với những ý kiến quan trọng: thêm dòng "Lời ghi nhận:" (in nghiêng) — lời
  cảm ơn ngắn gọn bằng tiếng Việt.
"""
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

OUT = (r"c:/Users/HLC/OneDrive/read_books/Lo-au/bai-bao-khgdvn/bài-03/"
       r"Giải trình phản hồi phản biện - Bài 2026.05.14.docx")
BLUE = RGBColor(0x00, 0x00, 0xCC)

doc = Document()
st = doc.styles["Normal"]
st.font.name = "Times New Roman"
st.font.size = Pt(13)


def H(t, sz=13):
    p = doc.add_paragraph(); r = p.add_run(t)
    r.bold = True; r.font.size = Pt(sz)
    return p


def body(t):
    p = doc.add_paragraph(); p.add_run(t); return p


def review(label, gop, phan, ack=None):
    """Một mục: số + góp ý phản biện (in nghiêng đậm) + Phản hồi (đậm + xanh) +
    Lời ghi nhận (in nghiêng) nếu là ý kiến quan trọng."""
    p = doc.add_paragraph()
    r0 = p.add_run(label + " "); r0.bold = True
    r1 = p.add_run("Góp ý của phản biện: "); r1.bold = True; r1.italic = True
    p.add_run(gop)
    p2 = doc.add_paragraph()
    r2 = p2.add_run("Phản hồi: "); r2.bold = True
    r3 = p2.add_run(phan); r3.font.color.rgb = BLUE
    if ack:
        p3 = doc.add_paragraph()
        r4 = p3.add_run("Lời ghi nhận: "); r4.bold = True; r4.italic = True
        r5 = p3.add_run(ack); r5.italic = True


# tiêu đề
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
tr = title.add_run("GIẢI TRÌNH PHẢN HỒI Ý KIẾN PHẢN BIỆN")
tr.bold = True; tr.font.size = Pt(15)
sub = doc.add_paragraph(); sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sr = sub.add_run("Bài báo mã 2026.05.14 — Môi trường gia đình và rối loạn "
                 "lo âu ở học sinh trung học cơ sở")
sr.bold = True; sr.italic = True; sr.font.size = Pt(13)
doc.add_paragraph()

body("Kính gửi Quý phản biện và Ban Biên tập Tạp chí Khoa học Giáo dục Việt Nam,")
body("Tác giả xin trân trọng cảm ơn Quý phản biện đã đọc kỹ bản thảo và đưa ra "
     "các góp ý cụ thể, có tính xây dựng cao. Tác giả đã rà soát và chỉnh sửa "
     "toàn bộ bản thảo theo từng điểm góp ý. Mọi chỗ sửa trong bản v8 đính kèm "
     "đều được TÔ ĐỎ để Quý phản biện dễ đối chiếu. Dưới đây là giải trình chi "
     "tiết từng điểm:")
doc.add_paragraph()

H("I. ĐIỂM CỐT LÕI XUYÊN SUỐT — đồng bộ thuật ngữ", 14)
body("Phản biện nhận xét (mục 2.3 phiếu + comment #0, #2, #5, #9): tiêu đề "
     "tiếng Việt không có cụm \"tổng quát\", trong khi tiêu đề tiếng Anh có "
     "\"GENERALIZED\" và một số đoạn thân bài dùng \"rối loạn lo âu tổng "
     "quát\". Quý phản biện ghi rõ: \"Trong nội dung bài viết, chủ yếu tập "
     "trung nói về Rối loạn lo âu\".")
review("→",
       "Đồng bộ thuật ngữ \"rối loạn lo âu\" xuyên suốt bài (cả tiếng Việt và "
       "tiếng Anh), giữ tên \"Generalized Anxiety Disorder\" / \"tổng quát\" "
       "duy nhất một chỗ ở mục Phương pháp nơi nêu tên gốc thang đo DSM-5 (vì "
       "đây là tên chính thức của thang đo). Hai chỗ còn lại có chữ \"tổng "
       "quát\" trong bài (đoạn 69 và 74) là tính từ với nghĩa khác — \"thang "
       "đo mang tính tổng quát\" và \"cảm nhận tổng quát\" — không liên quan "
       "đến tên rối loạn, nên được giữ nguyên.",
       "đồng bộ 12 vị trí trong bản v8 (tiêu đề EN, tóm tắt EN+VN, từ khóa "
       "EN+VN, các đoạn thân bài, caption Bảng 8/9, ghi chú).",
       ack="Tác giả trân trọng cảm ơn Quý phản biện đã chỉ ra điểm thiếu nhất "
           "quán quan trọng này. Việc đồng bộ thuật ngữ giúp tăng độ chính xác "
           "khoa học và độ rõ ràng xuyên suốt bản thảo.")
doc.add_paragraph()

H("II. GIẢI TRÌNH TỪNG ĐIỂM GÓP Ý", 14)
doc.add_paragraph()

review("1.",
       "Comment #0 — Tiêu đề tiếng Anh \"FAMILY ENVIRONMENT AND GENERALIZED "
       "ANXIETY DISORDER…\" không sát tiêu đề tiếng Việt vì tiếng Việt không "
       "có \"tổng quát\" (Generalized).",
       "đổi tiêu đề tiếng Anh thành \"FAMILY ENVIRONMENT AND ANXIETY DISORDER "
       "AMONG JUNIOR SECONDARY SCHOOL STUDENTS\" — đã bỏ \"GENERALIZED\". "
       "(Đoạn [0] trong v8, tô đỏ phần \"ANXIETY DISORDER\".)")

review("2.",
       "Mục 3.2 phiếu — Tóm tắt tiếng Anh cần chính xác và phù hợp hơn.",
       "tóm tắt tiếng Anh đã được rà lại; ba chỗ \"Generalized Anxiety Disorder "
       "(GAD)\" và \"GAD\" được thay bằng \"anxiety disorder\" / \"anxiety "
       "symptoms\" cho khớp tiêu đề và nội dung thân bài. (Đoạn [1] trong v8.)")

review("3.",
       "Comment #2 + mục 4.1, 4.2, 4.3 — Từ khóa cần chính xác (Rối loạn lo âu "
       "hay Rối loạn lo âu tổng quát?), hiệu đính lại tiếng Anh.",
       "đổi từ khóa VN từ \"rối loạn lo âu tổng quát\" sang \"rối loạn lo "
       "âu\"; từ khóa EN từ \"generalized anxiety disorder\" sang \"anxiety "
       "disorder\". Cả hai bộ từ khóa giữ đủ 5 từ trong khoảng yêu cầu 4–6. "
       "(Đoạn [2] và [9] trong v8.)")

review("4.",
       "Comment #4 + mục 5.5 phiếu — Tài liệu (Hiệp hội Tâm thần học Hoa Kỳ "
       "[American Psychiatric Association], 2013) cần chính xác tên, tác giả.",
       "chuẩn hóa định dạng trích dẫn theo APA 7: lần đầu ghi \"(American "
       "Psychiatric Association — APA, 2013)\". Các lần sau dùng \"(APA, "
       "2013)\" hoặc \"APA (2013)\". (Đoạn [24] trong v8.)")

review("5.",
       "Comment #5 — Đoạn nêu \"Thang đo rối loạn lo âu tổng quát\" cần thống "
       "nhất thuật ngữ.",
       "viết lại đoạn giới thiệu thang đo (Phương pháp): tên gốc của thang "
       "DSM-5 (Severity Measure for Generalized Anxiety Disorder, dành cho "
       "trẻ 11–17 tuổi; APA, 2013) được giữ tiếng Anh trong ngoặc đúng một "
       "lần để người đọc tra ngược, nội dung tiếp theo dùng \"thang đo mức độ "
       "nặng rối loạn lo âu\" và \"lo âu\" để đồng bộ với phần còn lại của "
       "bài. (Đoạn [38] trong v8.)",
       ack="Tác giả cảm ơn Quý phản biện đã nhấn mạnh tầm quan trọng của việc "
           "nêu chính xác tên gốc thang đo. Chỉnh sửa này giúp người đọc tra "
           "ngược được tài liệu nguồn và đảm bảo tính minh bạch về công cụ "
           "đo lường.")

review("6.",
       "Comment #6 + mục 6.3 phiếu — Khai báo sử dụng AI hiện đang viết theo "
       "kiểu cam kết, cần điều chỉnh.",
       "viết lại phần khai báo AI sang văn phong khai báo (declaration), bỏ "
       "các động từ cam kết. Nội dung mới mô tả khách quan công cụ AI đã được "
       "dùng (rà soát hành văn, định dạng trích dẫn, sắp xếp TLTK), nêu rõ "
       "thiết kế – phân tích – diễn giải là phần việc của tác giả, và khai "
       "báo AI không được ghi nhận là tác giả. (Đoạn [47] trong v8.)",
       ack="Tác giả cảm ơn góp ý tinh tế về văn phong. Phần khai báo AI nay đã "
           "được điều chỉnh đúng tinh thần khai báo minh bạch của Tạp chí, "
           "thay vì kiểu cam kết cá nhân.")

review("7.",
       "Comment #7 + mục 7.5 phiếu — Đoạn diễn đạt thuộc Bảng 6 dùng cụm \"hệ "
       "quả thứ cấp của quá trình lo âu kéo dài\". Chỉ số tương quan không "
       "hướng tới chỉ ra quan hệ nhân quả.",
       "bỏ ngôn ngữ nhân quả. Câu mới: \"...gợi ý đây có thể là biểu hiện "
       "ngoại vi hơn là biểu hiện cốt lõi trong cấu trúc lo âu của mẫu nghiên "
       "cứu.\" (Đoạn [59] trong v8.)",
       ack="Tác giả chân thành cảm ơn nhận xét chuyên môn quan trọng này. Việc "
           "loại bỏ ngôn ngữ nhân quả khỏi diễn giải kết quả tương quan là "
           "cải thiện đáng giá về tính chặt chẽ phương pháp luận.")

review("8.",
       "Comment #10 + mục 9.3 phiếu — Phần Tuyên bố xung đột lợi ích không "
       "đưa vào bài báo như hiện tại.",
       "xóa đoạn này khỏi thân bài. Vị trí được thay bằng một ghi chú đỏ rõ "
       "ràng để Quý phản biện biết đoạn đã được xử lý: \"[ĐÃ XÓA THEO GÓP Ý "
       "PHẢN BIỆN: Tuyên bố xung đột lợi ích sẽ được khai báo qua biểu mẫu "
       "nộp bài riêng của Tạp chí, không in trong thân bài.]\". Tác giả sẽ "
       "khai báo xung đột lợi ích qua biểu mẫu nộp bài của Tạp chí. (Đoạn "
       "[75] trong v8.)")

review("9.",
       "Comment #11, #12 + mục 10.1, 10.2 phiếu — Bổ sung tài liệu nghiên cứu "
       "ở Việt Nam (hiện có khá nhiều); sắp xếp lại Tài liệu tham khảo theo "
       "đúng quy chuẩn của Tạp chí (APA, thứ tự bảng chữ cái, đối với tác giả "
       "người Việt phải ghi đầy đủ họ tên).",
       "bổ sung 3 tài liệu tiếng Việt mới (đã được xác minh trực tiếp trên "
       "PubMed/PMC/nhà xuất bản): (a) Hoàng Trung Học & Nguyễn Thùy Dung "
       "(2025) trên American Journal of Psychiatric Rehabilitation; (b) Phạm "
       "Thị Thu Hoa và cộng sự (2024) trên Frontiers in Public Health; (c) "
       "Đinh Thị Hồng Vân, Đỗ Thị Lê Hằng & Phan Thị Mai Hương (2021) trên "
       "Psychology and Education. Cả ba được chèn vào danh mục theo thứ tự "
       "bảng chữ cái (foreign theo họ tác giả, người Việt theo tên) và ghi "
       "đầy đủ họ tên người Việt. Tổng TLTK hiện 24 mục (không vượt giới hạn "
       "25). (Cuối phần TÀI LIỆU THAM KHẢO trong v8, tô đỏ.)",
       ack="Tác giả cảm ơn Quý phản biện đã nhắc tới mảng nghiên cứu trong "
           "nước đang phát triển mạnh. Việc bổ sung các nghiên cứu Việt Nam "
           "mới giúp bản thảo đặt vào đúng bối cảnh học thuật trong nước "
           "hiện nay.")

review("10.",
       "Mục 2.2 phiếu — Tên bài tiếng Anh: Vừa phải (cần chuẩn hóa).",
       "tên bài tiếng Anh được rút gọn và làm sát tiếng Việt (xem mục 1 ở trên).")
doc.add_paragraph()

H("III. CÁC ĐIỂM PHẢN BIỆN ĐÃ ĐÁNH GIÁ \"TỐT\" — GIỮ NGUYÊN", 14)
body("Các nội dung được Quý phản biện cho điểm \"Tốt\" (chủ đề; tên bài tiếng "
     "Việt; tóm tắt tiếng Việt; đặt vấn đề; phương pháp; kết quả; thảo luận; "
     "kết luận; bảng biểu, sơ đồ và hình ảnh) được giữ nguyên về nội dung "
     "chính. Các điều chỉnh chỉ tập trung ở các điểm đã nêu trên và không làm "
     "thay đổi cấu trúc, dữ liệu hay phát hiện chính của nghiên cứu.")
doc.add_paragraph()

H("IV. TỆP ĐÍNH KÈM", 14)
body("• Bản thảo đã sửa: \"Công Thị Hằng_v8.docx\" — mọi chỗ sửa được tô đỏ "
     "để Quý phản biện đối chiếu nhanh.")
body("• Bản gốc đã gửi trước đó: \"Công Thị Hằng_v7.docx\" — giữ để đối "
     "chứng (không thay đổi).")
body("• Phiếu khảo sát (bản công cụ đo) — sẽ gửi kèm theo đúng yêu cầu của "
     "Tạp chí ở mục 6.")
doc.add_paragraph()

body("Một lần nữa, tác giả xin chân thành cảm ơn Quý phản biện và Ban Biên "
     "tập đã dành thời gian xem xét và đưa ra những góp ý quý báu giúp bản "
     "thảo được hoàn thiện hơn.")
doc.add_paragraph()
body("Trân trọng,")
body("Tác giả: Công Thị Hằng")

cp = doc.core_properties
cp.author = ""; cp.last_modified_by = ""
doc.save(OUT)
print("Đã lưu:", OUT)
