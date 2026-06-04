# -*- coding: utf-8 -*-
"""Đề cương nghiên cứu — CTH v5 — dựa trên 29+ bài NC (cập nhật 04/2026)"""
import sys, io
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

RED = RGBColor(0xFF, 0, 0); BLUE = RGBColor(0, 0x70, 0xC0)

doc = Document()
style = doc.styles['Normal']; style.font.name = 'Times New Roman'; style.font.size = Pt(13)
style.paragraph_format.space_after = Pt(6); style.paragraph_format.line_spacing = 1.5
for s in doc.sections: s.top_margin = Cm(2.5); s.bottom_margin = Cm(2.5); s.left_margin = Cm(3); s.right_margin = Cm(2)

def shade(cell, color):
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), color); s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)
def set_w(cell, w):
    tcW = cell._tc.get_or_add_tcPr(); we = OxmlElement('w:tcW'); we.set(qn('w:w'), str(int(w*567))); we.set(qn('w:type'), 'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t = doc.add_table(rows=1+len(rows), cols=len(headers)); t.style = 'Table Grid'; t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold = True; r.font.name = 'Times New Roman'; r.font.size = Pt(10)
        shade(c, 'D9E2F3')
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]; c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name = 'Times New Roman'; r.font.size = Pt(10)
def bl(text, bold=False, size=13):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold; r.font.color.rgb = BLUE
def nr(text, size=13, bold=False):
    p = doc.add_paragraph(); r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(size); r.bold = bold

# ============================================================
# TIÊU ĐỀ
# ============================================================
h = doc.add_heading('ĐỀ CƯƠNG NGHIÊN CỨU', 0)
for r in h.runs: r.font.name = 'Times New Roman'

h2 = doc.add_heading('LO ÂU Ở HỌC SINH TRUNG HỌC CƠ SỞ VÀ TRUNG HỌC PHỔ THÔNG TẠI VIỆT NAM:\nTỶ LỆ, YẾU TỐ LIÊN QUAN VÀ HIỆU QUẢ CAN THIỆP TẠI TRƯỜNG HỌC', 1)
for r in h2.runs: r.font.name = 'Times New Roman'

nr('Đề cương theo phong cách Công Thị Hằng v5. Dựa trên tổng hợp 29 bài nghiên cứu 2020\u20132025 (7 Việt Nam, 11 quốc tế ban đầu, 9 quốc tế bổ sung, 2 tổng quan toàn cầu).', size=11)
nr('Cập nhật: 04/2026.', size=11, bold=True)

# ============================================================
# 1. ĐẶT VẤN ĐỀ
# ============================================================
h1 = doc.add_heading('1. Đặt vấn đề', 1)
for r in h1.runs: r.font.name = 'Times New Roman'

# KT1: Phức tạp → đơn giản
nr('Rối loạn lo âu ở vị thành niên đã trở thành mối quan tâm sức khỏe công cộng toàn cầu trong thập kỷ qua. Nghiên cứu Gánh nặng Bệnh tật Toàn cầu 2021 cho thấy nhóm 10\u201314 tuổi tại các nước ASEAN chịu gánh nặng bệnh tật do rối loạn tâm thần cao nhất, chiếm 16,3% tổng DALYs ở nhóm tuổi này (GBD 2021 ASEAN, 2025). Nói cách khác, cứ 6 năm sống khỏe mạnh bị mất ở trẻ 10\u201314 tuổi ASEAN thì có 1 năm do rối loạn tâm thần.')

nr('Tại khu vực WHO châu Âu, 9 triệu thanh thiếu niên (10\u201319 tuổi) đang sống với rối loạn sức khỏe tâm thần, trong đó lo âu và trầm cảm chiếm hơn 50% tổng số ca (Lancet Regional Health Europe, 2025). Tại Hoa Kỳ, lo âu được chẩn đoán ở thanh thiếu niên đã tăng gấp đôi trong hệ thống y tế công lập — từ 9,6% (2013) lên 19,2% (2021), với AOR = 2,17 (JAACAP, 2024). Khảo sát quốc gia NHS 2025 tại Anh cho thấy 25,8% thanh niên 16\u201324 tuổi mắc rối loạn tâm thần phổ biến, tăng từ 18,9% năm 2014.')

# Bảng so sánh tỷ lệ
nr('Bảng 1. So sánh tỷ lệ lo âu ở học sinh THCS-THPT tại Việt Nam và quốc tế', bold=True, size=11)
tbl(
    ['Tác giả, năm', 'Quốc gia', 'n', 'Công cụ', 'Tỷ lệ lo âu'],
    [['V-NAMHS, 2022', 'Việt Nam (38 tỉnh)', '5.996', 'DISC-5/DSM-5', '2,3% (chẩn đoán)'],
     ['Vĩnh Lộc, 2024', 'Việt Nam (TPHCM)', '976', 'DASS-Y', '25,1%'],
     ['Hoa, 2024', 'Việt Nam (Hà Nội)', '3.910', 'GAD-7', '40,6%'],
     ['Hoàng Trung Học, 2025', 'Việt Nam (6 tỉnh)', '8.473', 'DASS-21', '41,5% (COVID)'],
     ['Danh Lâm, 2022', 'Việt Nam (Thanh Hóa)', '482', 'DASS-21', '49,0%'],
     ['Ngo Anh Vinh, 2024', 'Việt Nam (DTTS Lạng Sơn)', '845', 'DASS-21', '54,4%'],
     ['An Giang, 2025', 'Việt Nam (ĐBSCL)', '366', 'DASS-21', '61,2%'],
     ['Thảo Vi, 2025', 'Việt Nam (Huế)', '685', 'DASS-21', '65,8%'],
     ['Bảo Quyên, 2025', 'Việt Nam (Hà Nội)', '501', 'DASS-21', '86,2%'],
     ['Xu, 2021', 'Trung Quốc', '373.216', 'GAD-7', '9,89%'],
     ['Chen, 2023', 'Trung Quốc (Tây)', '63.205', 'GAD-7', '13,9%'],
     ['NSCH, 2020', 'Hoa Kỳ', '55.162', 'Chẩn đoán', '16,1%'],
     ['JAACAP, 2024', 'Hoa Kỳ', '13,7 triệu', 'Chẩn đoán', '19,2% (2021)']],
    widths=[3.5, 3.0, 2.0, 2.0, 2.5])
doc.add_paragraph()

# KT3: 3 bước xác nhận
nr('Tổng quan tài liệu của chúng tôi cho thấy ba khoảng trống quan trọng. Thứ nhất, tỷ lệ lo âu tại Việt Nam dao động từ 2,3% đến 86,2% — sự chênh lệch 37 lần phản ánh khoảng trống phương pháp luận nghiêm trọng: DISC-5 (công cụ chẩn đoán DSM-5) cho 2,3%, trong khi DASS-21 (sàng lọc với ngưỡng thấp) cho 62\u201386%. Nói cách khác, công cụ đo lường quyết định mạnh mẽ đến kết quả — nhưng chưa có nghiên cứu nào tại Việt Nam so sánh nhiều công cụ trên cùng một mẫu.')

nr('Thứ hai, chưa có thử nghiệm ngẫu nhiên có đối chứng (RCT) nào về can thiệp sức khỏe tâm thần tại trường học Việt Nam. Tổng quan hệ thống của Zhameden và cộng sự (2025) chỉ tìm được 6 RCTs tại tất cả các nước thu nhập thấp và trung bình, không có bài nào từ Việt Nam. Phân tích tổng hợp mạng Bayesian (BMC Psychiatry, 2025) trên 30 RCTs cho thấy CBT cá nhân hiệu quả nhất, nhưng bằng chứng chủ yếu từ phương Tây.')

nr('Thứ ba, nhóm dân tộc thiểu số (lo âu 54,4%, ACEs 48,9% — Ngô Anh Vinh, 2024) và nhóm nông thôn (lo âu nông thôn 11,33% > thành thị 8,77% — Xu, 2021) hầu như chưa được nghiên cứu. Vai trò trung gian của giấc ngủ (OR = 6,99\u201313,71), mạng xã hội (Nature Human Behaviour, 2025: VTN có rối loạn SKTT dùng MXH nhiều hơn) và áp lực học tập (OR = 11,58 — Wen, 2020) chưa được đánh giá trong mô hình nhân quả tại Việt Nam.')

# ============================================================
# 2. MỤC TIÊU
# ============================================================
h1 = doc.add_heading('2. Mục tiêu nghiên cứu', 1)
for r in h1.runs: r.font.name = 'Times New Roman'

nr('Mục tiêu tổng quát: Xác định tỷ lệ, yếu tố liên quan và đánh giá hiệu quả can thiệp tại trường học đối với lo âu ở học sinh THCS và THPT tại Việt Nam.', bold=True)

nr('Mục tiêu cụ thể:')
nr('(1) Xác định tỷ lệ rối loạn lo âu (chẩn đoán DISC-5/DSM-5) và triệu chứng lo âu (sàng lọc GAD-7) ở học sinh THCS và THPT tại 3 vùng sinh thái (đồng bằng, miền núi DTTS, đô thị). So sánh kết quả hai công cụ trên cùng mẫu để xác định ngưỡng cắt phù hợp cho quần thể Việt.')
nr('(2) Phân tích các yếu tố nguy cơ và bảo vệ: áp lực học tập (ESSA), mối quan hệ gia đình, sử dụng mạng xã hội, chất lượng giấc ngủ (PSQI), trải nghiệm bất lợi thời thơ ấu (ACEs), sự lạc quan (LOT-R), và hỗ trợ xã hội. Đánh giá vai trò trung gian của lo âu trong mối quan hệ lạc quan \u2192 trầm cảm.')
nr('(3) Thiết kế, triển khai và đánh giá hiệu quả một chương trình can thiệp sức khỏe tâm thần tại trường học (school-based intervention) sử dụng CBT nhóm — RCT đầu tiên tại Việt Nam.')

# ============================================================
# 3. PHƯƠNG PHÁP
# ============================================================
h1 = doc.add_heading('3. Phương pháp nghiên cứu', 1)
for r in h1.runs: r.font.name = 'Times New Roman'

nr('Nghiên cứu của chúng tôi sử dụng thiết kế hỗn hợp tuần tự giải thích (sequential explanatory mixed-methods design) gồm ba giai đoạn. Nói cách khác, nghiên cứu kết hợp khảo sát cắt ngang quy mô lớn, thử nghiệm can thiệp RCT và phỏng vấn sâu.')

h2 = doc.add_heading('3.1 Giai đoạn 1 — Khảo sát cắt ngang (năm 1)', 2)
for r in h2.runs: r.font.name = 'Times New Roman'
nr('Cỡ mẫu: 3.000\u20135.000 học sinh THCS và THPT, phân tầng theo vùng (đồng bằng, miền núi DTTS, đô thị) và cấp học (THCS/THPT). Lấy mẫu cụm nhiều bậc — chọn tỉnh → trường → lớp → toàn bộ HS trong lớp.')

nr('Bảng 2. Công cụ đo lường', bold=True, size=11)
tbl(
    ['Công cụ', 'Mục đích', 'Cơ sở'],
    [['GAD-7', 'Sàng lọc lo âu (ngưỡng ≥10)', 'Spitzer 2006; Hoa 2024 \u03b1=0,916'],
     ['DISC-5', 'Chẩn đoán rối loạn lo âu DSM-5', 'V-NAMHS 2022 — chuẩn vàng'],
     ['ESSA', 'Đo áp lực học tập', 'Vĩnh Lộc 2024 \u03b1=0,81'],
     ['PSQI', 'Chất lượng giấc ngủ', 'Chen 2023 OR=6,99'],
     ['ACEs', 'Trải nghiệm bất lợi thời thơ ấu', 'Ngo Anh Vinh 2024'],
     ['LOT-R', 'Sự lạc quan', 'Thảo Vi 2025 \u03b1=0,82'],
     ['Oslo-3', 'Hỗ trợ xã hội', 'Nakie 2022']],
    widths=[2.5, 4.5, 5.0])
doc.add_paragraph()

nr('Phân tích: Hồi quy logistic đa biến (AOR), phân tích hồ sơ tiềm ẩn (LPA) theo Qiu (2022) và Wen (2020), phân tích trung gian PROCESS macro theo Thảo Vi (2025). Mức ý nghĩa P < 0,05.')

h2 = doc.add_heading('3.2 Giai đoạn 2 — Thử nghiệm can thiệp RCT (năm 2)', 2)
for r in h2.runs: r.font.name = 'Times New Roman'
nr('Thiết kế: RCT cụm (cluster RCT), phân ngẫu nhiên theo trường. Nhóm can thiệp: CBT nhóm 8\u201312 buổi (dựa trên AJP, 2024: CBT phục hồi 47\u201366%). Nhóm đối chứng: chương trình giáo dục sức khỏe thông thường. Đánh giá trước\u2013sau\u20136 tháng.')
nr('Cơ sở chọn CBT: Phân tích tổng hợp mạng Bayesian (BMC Psychiatry, 2025) trên 30 RCTs xác nhận CBT cá nhân hiệu quả nhất, CBT nhóm xếp thứ 2\u20133. Tuy nhiên, Zhameden (2025) cảnh báo CBT chỉ hiệu quả 1/4 cho lo âu tại LMIC — cần điều chỉnh cho bối cảnh Việt Nam.')

h2 = doc.add_heading('3.3 Giai đoạn 3 — Phỏng vấn sâu (năm 2\u20133)', 2)
for r in h2.runs: r.font.name = 'Times New Roman'
nr('Phỏng vấn sâu 30\u201340 học sinh (chọn mẫu có mục đích từ GĐ1 và GĐ2) để hiểu trải nghiệm chủ quan, rào cản tiếp cận dịch vụ, và đánh giá quá trình can thiệp. Phân tích mã hóa chủ đề (thematic coding) theo mô hình Jenkins (2023).')

# ============================================================
# 4. Ý NGHĨA
# ============================================================
h1 = doc.add_heading('4. Ý nghĩa khoa học và thực tiễn', 1)
for r in h1.runs: r.font.name = 'Times New Roman'

# Kiến trúc B
nr('Dữ liệu tổng hợp từ 29 bài nghiên cứu, cho thấy lo âu ở thanh thiếu niên đang gia tăng toàn cầu (JAACAP 2024: tăng gấp đôi; NSCH 2020: tăng 61%; Na Uy: tăng liên tục 13 năm) trong khi Việt Nam chỉ có 8,4% thanh thiếu niên tiếp cận dịch vụ SKTT và cha mẹ chỉ nhận ra 5,1% con có vấn đề, gợi ý rằng nghiên cứu này là cấp thiết.', bold=True)

nr('Về khoa học, nghiên cứu sẽ là công trình đầu tiên tại Việt Nam: (a) sử dụng đồng thời DISC-5 và GAD-7 trên cùng một mẫu — xác định ngưỡng cắt phù hợp cho quần thể Việt; (b) thực hiện RCT can thiệp CBT tại trường học — lấp đầy khoảng trống 0 RCT mà Zhameden (2025) đã chỉ ra; (c) khám phá cơ chế trung gian lạc quan → lo âu → trầm cảm (Thảo Vi, 2025: β = \u20130,24) và vai trò mạng xã hội (Nature, 2025: VTN có rối loạn SKTT dùng MXH nhiều hơn).')

nr('Về thực tiễn, kết quả sẽ cung cấp dữ liệu cho Bộ Giáo dục và Đào tạo xây dựng chương trình SKTT học đường chuẩn hóa, đặc biệt cho nhóm DTTS (lo âu 54,4%, ACEs 48,9%) và vùng nông thôn — nơi tỷ lệ cao nhưng tiếp cận dịch vụ thấp nhất. Ba yếu tố CÓ THỂ CAN THIỆP ĐƯỢC — giấc ngủ (OR = 6,99\u201313,71), áp lực học tập (OR = 11,58), và chăm sóc cảm xúc gia đình (beta = \u20130,40) — sẽ là mục tiêu cụ thể cho chương trình can thiệp.')

# ============================================================
# 5. TIẾN ĐỘ
# ============================================================
h1 = doc.add_heading('5. Tiến độ dự kiến', 1)
for r in h1.runs: r.font.name = 'Times New Roman'

tbl(
    ['Giai đoạn', 'Thời gian', 'Nội dung chính'],
    [['Chuẩn bị', 'Tháng 1\u20136 năm 1', 'Xin đạo đức, chọn mẫu, đào tạo điều tra viên, pilot test'],
     ['GĐ1: Khảo sát', 'Tháng 7\u201312 năm 1', 'Thu thập dữ liệu 3.000\u20135.000 HS, GAD-7 + DISC-5'],
     ['Phân tích GĐ1', 'Tháng 1\u20134 năm 2', 'Hồi quy đa biến, LPA, PROCESS, xuất bản bài 1'],
     ['GĐ2: RCT', 'Tháng 5\u201312 năm 2', 'CBT nhóm 8\u201312 buổi, đánh giá trước\u2013sau'],
     ['GĐ3: Phỏng vấn', 'Tháng 1\u20136 năm 3', 'Phỏng vấn sâu 30\u201340 HS, phân tích chủ đề'],
     ['Follow-up', 'Tháng 7\u20139 năm 3', 'Đánh giá 6 tháng sau RCT'],
     ['Tổng hợp', 'Tháng 10\u201312 năm 3', 'Viết báo cáo, xuất bản bài 2\u20133, bảo vệ']],
    widths=[3.0, 3.5, 6.0])
doc.add_paragraph()

# ============================================================
# 6. TÀI LIỆU THAM KHẢO
# ============================================================
h1 = doc.add_heading('6. Tài liệu tham khảo (trích)', 1)
for r in h1.runs: r.font.name = 'Times New Roman'

refs = [
    '1. GBD 2021 ASEAN Mental Disorders Collaborators (2025). The Lancet Public Health, 10, e480.',
    '2. Hoa PTT et al. (2024). Frontiers in Public Health, 12, 1232856. [VN, GAD-7, n=3.910]',
    '3. Hoàng Trung Học & Nguyễn Thùy Dung (2025). Am J Psychiatric Rehabilitation, 28(1). [VN, DASS-21, n=8.473]',
    '4. Ngô Anh Vinh et al. (2024). J Affective Disorders Reports, 17. [VN DTTS, DASS-21+ACEs, n=845]',
    '5. V-NAMHS (2022). Viet Nam Adolescent Mental Health Survey. [VN, DISC-5, n=5.996]',
    '6. Zhameden et al. (2025). PLoS One, 20(4), e0316825. [6 RCTs LMIC, 0 VN]',
    '7. Xu Q et al. (2021). J Affective Disorders, 288, 17\u201322. [TQ, GAD-7, n=373.216]',
    '8. Chen Z et al. (2023). BMC Psychiatry, 23, 580. [TQ, n=63.205, giấc ngủ OR=6,99]',
    '9. Wen X et al. (2020). Int J Environ Res Public Health, 17(11), 4079. [TQ, LPA, áp lực OR=11,58]',
    '10. Hồ Thị Trúc Quỳnh (2025). Tạp chí Tâm lý học, 3(312). [VN Huế, PROCESS, lạc quan\u2192lo âu\u2192trầm cảm]',
    '11. Trần Hồ Vĩnh Lộc et al. (2024). TC Y học TPHCM, 27(5). [VN, DASS-Y, n=976]',
    '12. Nguyễn Ngọc Bảo Quyên et al. (2025). TC Y học Cộng đồng, 66. [VN Hà Nội, DASS-21, n=501]',
    '13. Nguyễn Danh Lâm et al. (2022). TC Y học Việt Nam, 516(1). [VN Thanh Hóa, DASS-21, n=482]',
    '14. JAACAP (2024). Trends in Mental Disorders. [Mỹ, n=13,7 triệu, lo âu 9,6%\u219219,2%]',
    '15. Lancet Regional Health Europe (2025). WHO Europe MH children/youth. [9 triệu VTN]',
    '16. Nature Human Behaviour (2025). Social media use in adolescents. [UK, n=3.340]',
    '17. AJP (2024). Pediatric Anxiety Treatment. [Tổng quan CBT 47\u201366% phục hồi]',
    '18. BMC Psychiatry (2025). CBT network meta-analysis. [30 RCTs, 1.711 trẻ]',
    '19. Pham et al. (2024). [VN Huế, chăm sóc cảm xúc beta=\u20130,40]',
    '20. Qiu Z et al. (2022). Frontiers in Psychology, 13. [TQ, nuôi dạy tích cực OR=0,30]',
]

for ref in refs:
    nr(ref, size=10)

doc.add_paragraph()
nr('Đề cương viết theo thuật toán bắt chước phong cách Công Thị Hằng v5. Dựa trên tổng hợp 29 bài nghiên cứu.', size=10, bold=True)

doc.save('De_cuong_Lo_au_CTH_v5_042026.docx')
sys.stderr.write('De_cuong_Lo_au_CTH_v5_042026.docx OK\n')
