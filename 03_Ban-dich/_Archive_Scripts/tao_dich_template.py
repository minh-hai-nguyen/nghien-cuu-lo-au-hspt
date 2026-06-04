# -*- coding: utf-8 -*-
"""
Template helper functions for creating translated papers
Following Jenkins et al. format:
- QR + link at top
- Page references from original journal
- Table Grid with borders
- Images at correct positions
- Red text for critique
- Gap analysis
"""
import os, sys, io
try:
    if hasattr(sys.stdout, 'buffer'):
        sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except:
    pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

IMAGES_DIR = os.path.join(os.path.dirname(os.path.abspath(__file__)), '..', 'papers', 'extracted_images')

def create_doc():
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    style.paragraph_format.space_after = Pt(6)
    style.paragraph_format.line_spacing = 1.5
    for s in doc.sections:
        s.top_margin = Cm(2.5)
        s.bottom_margin = Cm(2.5)
        s.left_margin = Cm(3)
        s.right_margin = Cm(2)
    return doc

def add_link_and_qr(doc, url, qr_filename):
    """Add link + QR code at top of document"""
    p = doc.add_paragraph()
    r = p.add_run(f'Link bài báo gốc: {url}')
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.font.color.rgb = RGBColor(0, 0, 255)

    # Add QR image
    qr_path = os.path.join(IMAGES_DIR, qr_filename)
    if os.path.exists(qr_path):
        p2 = doc.add_paragraph()
        p2.add_run().add_picture(qr_path, width=Cm(3))
        p3 = doc.add_paragraph()
        r3 = p3.add_run('Quét QR để truy cập bài báo gốc')
        r3.font.name = 'Times New Roman'
        r3.font.size = Pt(9)
        r3.italic = True

    doc.add_paragraph()

def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'
        r.font.color.rgb = RGBColor(0, 0, 0)
    return h

def add_page_ref(doc, page_num, journal_short, vol_info=''):
    """Add page reference line: --- Trang X, Journal, Vol ---"""
    p = doc.add_paragraph()
    info = f'--- Trang {page_num}, {journal_short}'
    if vol_info:
        info += f', {vol_info}'
    info += ' ---'
    r = p.add_run(info)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(9)
    r.italic = True
    r.font.color.rgb = RGBColor(128, 128, 128)

def add_p(doc, text, size=12, bold=False, italic=False):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.bold = bold
    r.italic = italic
    return p

def add_red(doc, text, size=12, bold=False):
    """Add red text (for critique/反biện)"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(size)
    r.font.color.rgb = RGBColor(255, 0, 0)
    r.bold = bold
    return p

def add_red_heading(doc, text):
    """Red bold heading for critique section"""
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(255, 0, 0)
    r.bold = True
    return p

def shade_cell(cell, color):
    s = OxmlElement('w:shd')
    s.set(qn('w:fill'), color)
    s.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(s)

def set_col_width(cell, width_cm):
    tcW = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW')
    w.set(qn('w:w'), str(int(width_cm * 567)))
    w.set(qn('w:type'), 'dxa')
    tcW.append(w)

def add_table(doc, headers, rows, widths=None):
    """Add Table Grid table with borders and proper column widths"""
    t = doc.add_table(rows=1+len(rows), cols=len(headers))
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Set widths
    if widths:
        for row in t.rows:
            for ci in range(len(headers)):
                set_col_width(row.cells[ci], widths[ci])

    # Header row
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade_cell(c, 'D9E2F3')

    # Data rows
    for ri, rd in enumerate(rows):
        for ci, v in enumerate(rd):
            c = t.rows[ri+1].cells[ci]
            c.text = str(v)
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)
    return t

def add_info_table(doc, data):
    """2-column info table (Mục | Chi tiết)"""
    t = doc.add_table(rows=len(data), cols=2)
    t.style = 'Table Grid'
    t.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        set_col_width(row.cells[0], 3.5)
        set_col_width(row.cells[1], 12.5)
    for i, (k, v) in enumerate(data):
        c0 = t.rows[i].cells[0]
        c0.text = k
        for p in c0.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade_cell(c0, 'D9E2F3')
        c1 = t.rows[i].cells[1]
        c1.text = v
        for p in c1.paragraphs:
            for r in p.runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
    return t

def add_image(doc, filename, width_cm=14, caption=''):
    """Add image from extracted_images folder"""
    img_path = os.path.join(IMAGES_DIR, filename)
    if os.path.exists(img_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        try:
            p.add_run().add_picture(img_path, width=Cm(width_cm))
        except Exception:
            p.add_run("[Hình không thể chèn — xem bản gốc PDF]")
        if caption:
            pc = doc.add_paragraph()
            pc.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = pc.add_run(caption)
            r.font.name = 'Times New Roman'
            r.font.size = Pt(9)
            r.italic = True
    else:
        add_p(doc, f'[Hình: {filename} — không tìm thấy file]', size=10, italic=True)

def add_abbreviation_table(doc, abbrevs):
    """Table of abbreviations"""
    add_heading(doc, 'BẢNG VIẾT TẮT', 2)
    t = doc.add_table(rows=1+len(abbrevs), cols=2)
    t.style = 'Table Grid'
    for row in t.rows:
        set_col_width(row.cells[0], 3.0)
        set_col_width(row.cells[1], 13.0)
    # Header
    for i, h in enumerate(['Viết tắt', 'Đầy đủ']):
        c = t.rows[0].cells[i]
        c.text = h
        for p in c.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.name = 'Times New Roman'
                r.font.size = Pt(10)
        shade_cell(c, 'D9E2F3')
    # Data
    for ri, (abbr, full) in enumerate(abbrevs):
        t.rows[ri+1].cells[0].text = abbr
        t.rows[ri+1].cells[1].text = full
        for ci in range(2):
            for p in t.rows[ri+1].cells[ci].paragraphs:
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(10)

try:
    print('Template module loaded OK')
except:
    pass
