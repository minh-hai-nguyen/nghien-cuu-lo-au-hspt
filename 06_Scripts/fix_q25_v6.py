# -*- coding: utf-8 -*-
"""
Q2.5 v5 -> v6: sửa các điểm phát hiện qua re-check (3 agent, 17/05/2026).
- Bổ sung trích dẫn (Olweus, 1996) và (Rosenberg, 1965) ở lần nhắc đầu trong
  Phương pháp — trước đây hai mục này là tài liệu mồ côi (có trong TLTK, không
  được trích trong thân bài).
- Làm mềm ngôn ngữ nhân-quả trong Bảng 1 phụ lục bản VN ("làm tăng"/"tăng" ->
  "liên hệ thuận với") để nhất quán với cảnh báo SEM cắt ngang đã thêm ở v5.
"""
import io, sys
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding="utf-8")
from docx import Document


def rep_para(doc, idx, old, new, label):
    p = doc.paragraphs[idx]
    for r in p.runs:
        if old in r.text:
            r.text = r.text.replace(old, new)
            print(f"  [{idx}] OK: {label}")
            return True
    if old in p.text and p.runs:
        p.runs[0].text = p.text.replace(old, new)
        for r in p.runs[1:]:
            r.text = ""
        print(f"  [{idx}] OK (gộp run): {label}")
        return True
    print(f"  [{idx}] !! KHÔNG khớp: {label}")
    return False


def rep_cell(cell, old, new):
    for p in cell.paragraphs:
        for r in p.runs:
            if old in r.text:
                r.text = r.text.replace(old, new)
                return True
        if old in p.text and p.runs:
            p.runs[0].text = p.text.replace(old, new)
            for r in p.runs[1:]:
                r.text = ""
            return True
    return False


# ===== EN v5 -> v6 =====
print("=== Q25 EN v5 -> v6 ===")
en = Document("bai-bao-khgdvn/Q25_SEM_Pathways_EN_v5.docx")
rep_para(en, 15,
    "school bullying with an adapted Olweus Bully/Victim instrument covering physical, verbal, and cyber bullying, self-esteem with the ten-item Rosenberg Self-Esteem Scale,",
    "school bullying with an adapted Olweus Bully/Victim instrument (Olweus, 1996) covering physical, verbal, and cyber bullying, self-esteem with the ten-item Rosenberg Self-Esteem Scale (Rosenberg, 1965),",
    "thêm trích dẫn Olweus (1996) + Rosenberg (1965)")
en.core_properties.author = ""
en.core_properties.last_modified_by = ""
en.save("bai-bao-khgdvn/Q25_SEM_Pathways_EN_v6.docx")
print("  -> đã lưu Q25_SEM_Pathways_EN_v6.docx\n")

# ===== VN v5 -> v6 =====
print("=== Q25 VN v5 -> v6 ===")
vn = Document("bai-bao-khgdvn/Q25_SEM_Pathways_VN_v5.docx")
rep_para(vn, 14,
    "bắt nạt học đường bằng phiên bản phỏng theo Olweus Bully/Victim cho cả bắt nạt thể chất, lời nói và mạng, lòng tự trọng bằng thang Rosenberg 10 mục,",
    "bắt nạt học đường bằng phiên bản phỏng theo Olweus Bully/Victim (Olweus, 1996) cho cả bắt nạt thể chất, lời nói và mạng, lòng tự trọng bằng thang Rosenberg 10 mục (Rosenberg, 1965),",
    "thêm trích dẫn Olweus (1996) + Rosenberg (1965)")

# Bảng 1 phụ lục VN — làm mềm ngôn ngữ nhân quả (cột Diễn giải)
CELLFIX = [
    (1,  "làm tăng mạnh lo âu tổng quát", "liên hệ thuận mạnh với lo âu tổng quát"),
    (2,  "làm tăng mạnh lo âu tổng quát", "liên hệ thuận mạnh với lo âu tổng quát"),
    (3,  "làm tăng mạnh lo âu xã hội",    "liên hệ thuận mạnh với lo âu xã hội"),
    (6,  "làm tăng lo âu xã hội",         "liên hệ thuận với lo âu xã hội"),
    (7,  "bắt nạt học đường tăng lo âu chia ly", "bắt nạt học đường liên hệ thuận với lo âu chia ly"),
    (8,  "nghiện điện thoại tăng lo âu tổng quát", "nghiện điện thoại liên hệ thuận với lo âu tổng quát"),
    (10, "áp lực học tập tăng lo âu chia ly", "áp lực học tập liên hệ thuận với lo âu chia ly"),
    (14, "KHÔNG có tác động lên lo âu chia ly", "KHÔNG có liên hệ trực tiếp với lo âu chia ly"),
]
# tìm Bảng 1 phụ lục (15 hàng, cột Đường dẫn|β|Diễn giải)
tb1 = None
for tb in vn.tables:
    if len(tb.rows) == 15 and tb.rows[0].cells[0].text.strip() == "Đường dẫn":
        tb1 = tb
        break
if tb1 is None:
    print("  !! không tìm thấy Bảng 1 phụ lục VN")
else:
    n = 0
    for ri, old, new in CELLFIX:
        if rep_cell(tb1.rows[ri].cells[2], old, new):
            n += 1
        else:
            print(f"  !! Bảng 1 R{ri}: không khớp '{old[:30]}'")
    print(f"  Bảng 1 phụ lục: làm mềm {n}/8 dòng ngôn ngữ nhân quả")

vn.core_properties.author = ""
vn.core_properties.last_modified_by = ""
vn.save("bai-bao-khgdvn/Q25_SEM_Pathways_VN_v6.docx")
print("  -> đã lưu Q25_SEM_Pathways_VN_v6.docx")
print("[DONE]")
