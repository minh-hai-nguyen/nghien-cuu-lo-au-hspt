# -*- coding: utf-8 -*-
"""Full pipeline:
1. Restore tu BEFORE_TOC
2. Apply heading style nhat quan (context-aware)
3. Fix cac heading style ngwn nhau
4. Extract pages tu Word COM
5. Update TOC table
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')
BACKUP_CLEAN = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026_BEFORE_TOC.docx')


# ============================================================
# STEP 1: Restore
# ============================================================
print("[STEP 1] Restoring from BEFORE_TOC...")
shutil.copy(BACKUP_CLEAN, FILE)


# ============================================================
# STEP 2: Apply heading styles
# ============================================================
print("\n[STEP 2] Apply heading styles (context-aware)...")
d = Document(FILE)

def in_table(p):
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag in (qn('w:tbl'), qn('w:tc')):
            return True
        parent = parent.getparent()
    return False

H1_EXACT = {
    'MỞ ĐẦU', 'KẾT LUẬN VÀ KIẾN NGHỊ',
    'TÀI LIỆU THAM KHẢO', 'PHỤ LỤC', 'MỤC LỤC',
    'LỜI CAM ĐOAN', 'LỜI CẢM ƠN',
    'DANH MỤC CÔNG TRÌNH ĐÃ CÔNG BỐ',
    'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
    'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
}
CHUONG_RE = re.compile(r'^CHƯƠNG\s+[\dIVX]+\s*[:.]?\s*$', re.IGNORECASE)
TIEU_KET_RE = re.compile(r'^TIỂU KẾT CHƯƠNG\s+[\dIVX]+\s*$', re.IGNORECASE)
CHAPTER_TITLE_KEYWORDS = ['CƠ SỞ LÝ LUẬN', 'TỔ CHỨC VÀ PHƯƠNG PHÁP',
                          'KẾT QUẢ NGHIÊN CỨU', 'PHÂN TÍCH KẾT QUẢ',
                          'ĐỀ XUẤT']

def count_numbers_in_prefix(text):
    m = re.match(r'^(\d+(?:\.\d+)*)\.?\s', text)
    if not m: return None
    return len(m.group(1).split('.'))

def determine_context(prev_h1_text):
    if not prev_h1_text: return 'other'
    upper = prev_h1_text.upper().strip()
    if 'MỞ ĐẦU' in upper or 'KẾT LUẬN' in upper:
        return 'mo_dau'
    if 'CHƯƠNG' in upper or any(kw in upper for kw in CHAPTER_TITLE_KEYWORDS) or 'TIỂU KẾT' in upper:
        return 'chuong'
    return 'other'

def decide_heading(text, context):
    text = text.strip()
    if not text or len(text) > 250: return None
    if re.match(r'^(Bảng|Hình|Biểu đồ|Sơ đồ)\s+\d', text, re.IGNORECASE): return None
    if re.match(r'^[a-z]/\s', text): return None
    upper = text.upper()
    if upper in H1_EXACT: return 1
    if CHUONG_RE.match(text): return 1
    if TIEU_KET_RE.match(text): return 1
    if upper == text and len(text) < 200 and any(kw in upper for kw in CHAPTER_TITLE_KEYWORDS):
        return 1
    if re.match(r'^Giả thuyết\s+[1-9]\s*[—–-]', text): return 3
    n = count_numbers_in_prefix(text)
    if n is None: return None
    if context == 'mo_dau':
        return min(n + 1, 5)
    elif context == 'chuong':
        if n < 2: return None
        return min(n, 5)
    return None


prev_h1 = ''
context = 'other'
applied = 0
demoted = 0  # heading -> normal
relevel = 0  # heading at wrong level -> correct level

for i, p in enumerate(d.paragraphs):
    if in_table(p): continue
    text = p.text.strip()
    if not text: continue
    if i < 60: continue

    current_style = p.style.name if p.style else 'Normal'

    # Update context tracking
    upper = text.upper()
    is_h1_text = (upper in H1_EXACT or CHUONG_RE.match(text) or TIEU_KET_RE.match(text)
                  or (upper == text and len(text) < 200 and any(kw in upper for kw in CHAPTER_TITLE_KEYWORDS)))
    if is_h1_text:
        prev_h1 = text
        context = determine_context(prev_h1)

    target_level = decide_heading(text, context)
    is_currently_heading = ('Heading' in current_style or 'Tiêu đề' in current_style)

    if target_level is None:
        # Should NOT be heading
        if is_currently_heading:
            # Currently styled as heading but my rule says NOT heading
            # DEMOTE to Normal (this fixes wrongly-applied Heading)
            try:
                p.style = d.styles['Normal']
                demoted += 1
            except: pass
        continue

    # Should be heading at target_level
    expected_styles = [f'Heading {target_level}', f'Tiêu đề {target_level}']
    if current_style in expected_styles: continue

    # Apply
    try:
        p.style = d.styles[f'Heading {target_level}']
        if is_currently_heading:
            relevel += 1
        else:
            applied += 1
    except: pass

print(f"  Apply Normal->Heading: {applied}")
print(f"  Re-level Heading: {relevel}")
print(f"  Demote Heading->Normal: {demoted}")

d.save(FILE)
print(f"  Saved")


# ============================================================
# STEP 3: Word COM - extract heading + page
# ============================================================
print("\n[STEP 3] Word COM extract...")
import win32com.client, pythoncom
TEMP = os.path.join(ROOT, 'Luận án TS', '_temp_pipeline.docx')
shutil.copy(FILE, TEMP)
pythoncom.CoInitialize()
word = None
headings = []
try:
    word = win32com.client.DispatchEx('Word.Application')
    word.Visible = False
    word.DisplayAlerts = False
    doc = word.Documents.Open(os.path.abspath(TEMP), ReadOnly=False)
    doc.Fields.Update()
    doc.Repaginate()
    print(f"  Pages: {doc.ComputeStatistics(2)}")
    for level in [1, 2, 3, 4, 5]:
        for style_name in [f'Heading {level}', f'Tiêu đề {level}']:
            try:
                rng = doc.Content
                find = rng.Find
                find.ClearFormatting()
                find.Style = doc.Styles(style_name)
                find.Text = ''
                find.Forward = True
                find.Wrap = 0
                find.Format = True
                while find.Execute():
                    para = rng.Paragraphs.First
                    text = para.Range.Text.strip().replace('\r','').replace('\x07','')
                    if para.Range.Tables.Count > 0:
                        rng.Collapse(0); continue
                    if not text or len(text) > 250:
                        rng.Collapse(0); continue
                    page = para.Range.Information(3)
                    headings.append({'level': level, 'text': text, 'page': page})
                    rng.Collapse(0)
            except: pass
    doc.Close(SaveChanges=False)
finally:
    if word:
        try: word.Quit()
        except: pass
    pythoncom.CoUninitialize()
    if os.path.exists(TEMP):
        try: os.remove(TEMP)
        except: pass

# Dedupe + sort
seen = set(); unique = []
for h in headings:
    k = (h['text'], h['page'])
    if k not in seen:
        seen.add(k); unique.append(h)
headings = unique
print(f"  Extracted {len(headings)} unique headings")


# ============================================================
# STEP 4: Update TOC table
# ============================================================
print("\n[STEP 4] Update TOC table...")
d = Document(FILE)
toc_table = None
for tb in d.tables:
    if len(tb.rows) > 0 and tb.rows[0].cells[0].text.strip().upper() == 'MỞ ĐẦU':
        toc_table = tb; break
if toc_table is None:
    print("KHONG TIM THAY TOC TABLE")
    sys.exit(1)

def normalize(s):
    s = s.strip().lower()
    s = re.sub(r'\s+', ' ', s)
    s = re.sub(r'^[\d\.]+\s*', '', s)
    return s.strip(' .')

lookup = {}
for h in headings:
    k = normalize(h['text'])
    if k and k not in lookup:
        lookup[k] = h['page']
    full = re.sub(r'\s+', ' ', h['text'].strip().lower())
    if full and full not in lookup:
        lookup[full] = h['page']

updated = 0
not_found = []
for ri, row in enumerate(toc_table.rows):
    name = row.cells[0].text.strip()
    old_page = row.cells[1].text.strip() if len(row.cells) > 1 else ''
    norm = normalize(name)
    full = re.sub(r'\s+', ' ', name.strip().lower())
    new_page = lookup.get(norm) or lookup.get(full)
    if new_page is None and norm:
        for k, v in lookup.items():
            if k and len(norm) > 8 and (norm in k or k in norm):
                new_page = v; break
    if new_page is None:
        not_found.append((ri, name)); continue
    new_page_str = str(new_page)
    if new_page_str != old_page:
        cell = row.cells[1]
        font_name = 'Times New Roman'; font_size = Pt(13)
        if cell.paragraphs and cell.paragraphs[0].runs:
            old_r = cell.paragraphs[0].runs[0]
            if old_r.font.name: font_name = old_r.font.name
            if old_r.font.size: font_size = old_r.font.size
        for p in cell.paragraphs:
            for r in p.runs:
                r.text = ''
        cell.paragraphs[0].text = ''
        new_run = cell.paragraphs[0].add_run(new_page_str)
        new_run.font.name = font_name
        new_run.font.size = font_size
        updated += 1

print(f"  Updated {updated} page numbers")
print(f"  Not found: {len(not_found)} rows")

d.save(FILE)
print(f"\n=== DONE === Saved: {FILE}")

# Show remaining unmatched
if not_found:
    print(f"\nCac muc TOC table KHONG MATCH heading nao trong doc:")
    for ri, name in not_found[:15]:
        print(f"  Row {ri}: {name}")
    if len(not_found) > 15:
        print(f"  ... va {len(not_found)-15} muc khac")

# Show extracted headings NOT in TOC table (might need adding)
print(f"\n=== Headings co trong doc nhung CHUA co trong TOC table ===")
toc_names = [normalize(row.cells[0].text.strip()) for row in toc_table.rows]
toc_full = [re.sub(r'\s+', ' ', row.cells[0].text.strip().lower()) for row in toc_table.rows]
missing = []
for h in headings:
    norm = normalize(h['text'])
    full = re.sub(r'\s+', ' ', h['text'].strip().lower())
    in_toc = (norm in toc_names) or (full in toc_full)
    if not in_toc:
        # Also try fuzzy
        in_toc = any(norm and (norm in t or (len(t) > 8 and t in norm)) for t in toc_names)
    if not in_toc:
        missing.append(h)

print(f"  {len(missing)} headings missing from TOC table:")
for h in missing[:20]:
    print(f"    L{h['level']} p{h['page']}: {h['text'][:80]}")
