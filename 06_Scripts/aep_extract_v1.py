# -*- coding: utf-8 -*-
"""
aep_extract_v1.py — Assertion-Evidence Pair extraction
Extract structured {assertion, evidence} from 68 summaries + translations.
Then cross-check vs report claims.

Evidence schema:
  source: canonical_id
  metric: OR|AOR|beta|d|g|SMD|SUCRA
  value: float
  CI: [lower, upper] or None
  p: str or None
  design: cross-sectional|cohort|RCT|meta-analysis|...
  outcome: Anxiety|Depression|Stress|SAD|Mixed
  instrument: GAD-7|DASS-21|PHQ-9|...
  n: int or None
  population: str
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
TT_DIR = os.path.join(ROOT, 'Tom-tat-tung-bai')
BD_DIR = os.path.join(ROOT, '03_Ban-dich')
REPORT = os.path.join(ROOT, '01_Bao-cao',
                      'Bao cao Can thiep tam ly RLLA VTN - 12042026 v5-4.docx')

# ============================================================
# Helpers
# ============================================================
INSTRUMENTS = {
    'GAD-7': 'Anxiety', 'DASS-21': 'Mixed', 'DASS-Y': 'Mixed',
    'DASS-42': 'Mixed', 'PHQ-9': 'Depression', 'CESD-R': 'Depression',
    'CES-D': 'Depression', 'MHT': 'Anxiety', 'STAI': 'Anxiety',
    'SAS': 'Anxiety', 'SIAS': 'SAD', 'PSQI': 'Sleep', 'ESSA': 'AcademicStress',
    'SDQ': 'General', 'SCAS': 'Anxiety', 'IGDS9-SF': 'Gaming',
    'MPVS': 'Bullying', 'CSES': 'CopingEfficacy',
}

DESIGN_MAP = {
    'meta-analysis': 'meta-analysis', 'NMA': 'meta-analysis',
    'network meta-analysis': 'meta-analysis', 'umbrella': 'umbrella-review',
    'systematic review': 'systematic-review', 'SR': 'systematic-review',
    'RCT': 'RCT', 'cluster RCT': 'RCT', 'cluster controlled': 'quasi-RCT',
    'randomized': 'RCT', 'randomised': 'RCT', 'thử nghiệm': 'RCT',
    'longitudinal': 'cohort', 'cohort': 'cohort', 'NC dọc': 'cohort',
    'cross-sectional': 'cross-sectional', 'cắt ngang': 'cross-sectional',
    'khảo sát': 'cross-sectional', 'scoping': 'scoping-review',
}

def read_docx(path):
    try:
        d = Document(path)
        lines = [p.text for p in d.paragraphs if p.text.strip()]
        for t in d.tables:
            for r in t.rows:
                lines.append(' | '.join(c.text.strip() for c in r.cells))
        return '\n'.join(lines)
    except:
        return ''

def get_cid(fn):
    m = re.match(r'(VN\d{3}|QT\d{3})', fn)
    return m.group(1) if m else None

def detect_design(text):
    tl = text.lower()
    for key, val in sorted(DESIGN_MAP.items(), key=lambda x: -len(x[0])):
        if key.lower() in tl:
            return val
    return 'cross-sectional'

def detect_instruments(text):
    found = []
    for inst in INSTRUMENTS:
        if inst in text:
            found.append(inst)
    return found

def detect_outcome(text):
    outcomes = set()
    tl = text.lower()
    if 'lo âu' in tl or 'anxiety' in tl: outcomes.add('Anxiety')
    if 'trầm cảm' in tl or 'depression' in tl: outcomes.add('Depression')
    if 'stress' in tl or 'căng thẳng' in tl: outcomes.add('Stress')
    if 'sad' in text or 'lo âu xã hội' in tl or 'social anxiety' in tl: outcomes.add('SAD')
    return list(outcomes) or ['Unknown']

def extract_sample_size(text):
    # Try n = X patterns
    patterns = [
        r'[Nn]\s*=\s*([\d.,]+)',
        r'(\d{3,6}(?:\.\d{3})*)\s*(?:học sinh|HS|VTN|sinh viên|người|bệnh nhân|hồ sơ)',
        r'mẫu\s+(\d{3,6})',
    ]
    for pat in patterns:
        m = re.search(pat, text)
        if m:
            val = m.group(1).replace('.', '').replace(',', '')
            try:
                return int(val)
            except:
                pass
    return None

def extract_population(text):
    """Extract short population description."""
    pats = [
        r'(\d+\s*(?:HS|học sinh|VTN|sinh viên).*?(?:tại|ở)\s+[^,.\n]{5,40})',
        r'((?:HS|học sinh|VTN).*?(?:THPT|THCS|trung học|đại học)[^,.\n]{0,30})',
    ]
    for pat in pats:
        m = re.search(pat, text[:500])
        if m:
            return m.group(1).strip()[:80]
    return ''

# ============================================================
# Main AEP extraction
# ============================================================
print('='*70)
print('AEP EXTRACTION v1 — Assertion-Evidence Pairs')
print('='*70)

# Load all files
print('\nLoading files...')
summaries, translations = {}, {}
for f in sorted(os.listdir(TT_DIR)):
    if f.endswith('.docx') and not f.startswith('~') and not f.startswith('_'):
        cid = get_cid(f)
        if cid: summaries[cid] = read_docx(os.path.join(TT_DIR, f))
for f in sorted(os.listdir(BD_DIR)):
    if f.endswith('.docx') and not f.startswith('~') and not f.startswith('_'):
        cid = get_cid(f)
        if cid: translations[cid] = read_docx(os.path.join(BD_DIR, f))
for subdir in os.listdir(BD_DIR):
    sp = os.path.join(BD_DIR, subdir)
    if os.path.isdir(sp):
        for f in os.listdir(sp):
            cid = get_cid(f)
            if cid and cid not in translations:
                if f.endswith('.docx'):
                    translations[cid] = read_docx(os.path.join(sp, f))
                elif f.endswith('.md'):
                    try:
                        with open(os.path.join(sp, f), encoding='utf-8') as fh:
                            translations[cid] = fh.read()
                    except: pass
print(f'  Summaries: {len(summaries)}, Translations: {len(translations)}')

# Extract AEPs
EFFECT_PAT = re.compile(
    r'(OR|AOR|aOR|β|beta|Cohen\s*d|d|Hedges?\s*g|g|SMD|SUCRA)\s*=\s*'
    r'([-−]?\d+[,.]?\d*)'
    r'(?:\s*\((?:KTC|CI|95\s*%)\s*[:\s]*'
    r'([-−]?\d+[,.]?\d*)\s*[-–]\s*([-−]?\d+[,.]?\d*)\))?'
    r'(?:\s*[,;]?\s*(?:p\s*[=<]\s*)([\d,.<]+))?'
)

all_aeps = {}
total_aeps = 0

for cid in sorted(set(summaries.keys()) | set(translations.keys())):
    sum_txt = summaries.get(cid, '')
    trans_txt = translations.get(cid, '')
    combined = sum_txt + '\n' + trans_txt

    design = detect_design(combined)
    instruments = detect_instruments(combined)
    outcomes = detect_outcome(combined)
    n = extract_sample_size(combined)
    population = extract_population(combined)

    paper_aeps = []

    # Extract from summary (preferred — curated)
    source_txt = sum_txt if sum_txt else trans_txt
    for m in EFFECT_PAT.finditer(source_txt):
        metric = m.group(1).strip()
        # Normalize metric name
        metric_norm = metric.replace('Cohen d', 'd').replace('Hedges g', 'g').replace('beta', 'β')
        if metric_norm.lower() in ('aor', 'or'): metric_norm = metric_norm.upper()

        value = m.group(2).replace(',', '.').replace('−', '-')
        ci_lo = m.group(3).replace(',', '.').replace('−', '-') if m.group(3) else None
        ci_hi = m.group(4).replace(',', '.').replace('−', '-') if m.group(4) else None
        p_val = m.group(5) if m.group(5) else None

        # Get context (assertion)
        start = max(0, m.start() - 60)
        end = min(len(source_txt), m.end() + 60)
        assertion = source_txt[start:end].replace('\n', ' ').strip()

        # Detect specific outcome for this effect size
        es_context = source_txt[max(0, m.start()-100):m.end()+100].lower()
        es_outcome = 'Unknown'
        if 'lo âu' in es_context or 'anxiety' in es_context: es_outcome = 'Anxiety'
        elif 'trầm cảm' in es_context or 'depression' in es_context: es_outcome = 'Depression'
        elif 'stress' in es_context or 'căng thẳng' in es_context: es_outcome = 'Stress'
        elif 'sad' in es_context or 'lo âu xã hội' in es_context: es_outcome = 'SAD'
        else: es_outcome = outcomes[0] if outcomes else 'Unknown'

        # Detect specific instrument context
        es_instrument = None
        for inst in instruments:
            if inst.lower() in es_context:
                es_instrument = inst
                break
        if not es_instrument and instruments:
            es_instrument = instruments[0]

        aep = {
            'assertion': assertion,
            'evidence': {
                'source': cid,
                'metric': metric_norm,
                'value': float(value) if value else None,
                'CI': [float(ci_lo), float(ci_hi)] if ci_lo and ci_hi else None,
                'p': p_val,
                'design': design,
                'outcome': es_outcome,
                'instrument': es_instrument,
                'n': n,
                'population': population,
            }
        }
        paper_aeps.append(aep)

    all_aeps[cid] = paper_aeps
    total_aeps += len(paper_aeps)

print(f'\nTotal AEPs extracted: {total_aeps}')
print(f'Papers with AEPs: {sum(1 for v in all_aeps.values() if v)}')

# ============================================================
# Cross-check AEPs vs Report
# ============================================================
print('\n' + '='*70)
print('CROSS-CHECK: Report claims vs AEPs')
print('='*70)

report_doc = Document(REPORT)
report_issues = []

for pi, p in enumerate(report_doc.paragraphs):
    ptxt = p.text.strip()
    if not ptxt:
        continue

    cids = re.findall(r'\[(VN\d{3}|QT\d{3})\]', ptxt)
    if not cids:
        continue

    # Find effect sizes in paragraph
    for m in EFFECT_PAT.finditer(ptxt):
        metric = m.group(1).strip()
        value_str = m.group(2).replace(',', '.').replace('−', '-')
        try:
            value = float(value_str)
        except:
            continue

        # Check against AEPs for cited papers
        for cid in cids:
            paper_aeps = all_aeps.get(cid, [])
            # Find matching AEP
            matched = False
            for aep in paper_aeps:
                ev = aep['evidence']
                if ev['value'] and abs(ev['value'] - value) < 0.02:
                    matched = True
                    # Check: is report using causal language with cross-sectional?
                    causal_words = ['gây ra', 'làm tăng', 'làm giảm', 'giảm nguy cơ',
                                    'tăng nguy cơ', 'hiệu quả can thiệp']
                    caveats = ['cắt ngang', 'cross-section', 'hiệp hội', 'association',
                               'lưu ý', 'caveat', 'ứng viên']
                    if ev['design'] == 'cross-sectional':
                        for cw in causal_words:
                            if cw in ptxt and not any(cv in ptxt for cv in caveats):
                                report_issues.append({
                                    'type': 'causal_with_crosssectional',
                                    'para': pi,
                                    'paper': cid,
                                    'metric': metric,
                                    'value': value,
                                    'design': ev['design'],
                                    'causal_word': cw,
                                    'context': ptxt[:120],
                                })
                                break

                    # Check: is report claiming wrong outcome?
                    if ev['outcome'] != 'Unknown':
                        report_ctx = ptxt.lower()
                        if ev['outcome'] == 'Depression' and 'lo âu' in report_ctx and 'trầm cảm' not in report_ctx:
                            report_issues.append({
                                'type': 'outcome_mismatch',
                                'para': pi,
                                'paper': cid,
                                'metric': metric,
                                'value': value,
                                'aep_outcome': ev['outcome'],
                                'report_context': ptxt[:120],
                                'note': f'AEP says outcome={ev["outcome"]} but report context implies Anxiety',
                            })
                    break

            if not matched:
                # Value cited with paper but no AEP matches → might be cross-ref
                pass

print(f'  Report issues from AEP cross-check: {len(report_issues)}')
for ri in report_issues[:10]:
    print(f'    P{ri["para"]} [{ri["paper"]}] {ri["type"]}: {ri.get("note", ri.get("causal_word", ""))}')

# ============================================================
# Statistics
# ============================================================
print('\n' + '='*70)
print('AEP STATISTICS')
print('='*70)

designs = {}
for cid, aeps in all_aeps.items():
    if aeps:
        d = aeps[0]['evidence']['design']
        designs[d] = designs.get(d, 0) + 1

print('Papers by design:')
for d, c in sorted(designs.items(), key=lambda x: -x[1]):
    print(f'  {d}: {c}')

has_ci = sum(1 for aeps in all_aeps.values() for a in aeps if a['evidence']['CI'])
has_p = sum(1 for aeps in all_aeps.values() for a in aeps if a['evidence']['p'])
print(f'\nAEPs with CI: {has_ci}/{total_aeps} ({100*has_ci//max(total_aeps,1)}%)')
print(f'AEPs with p-value: {has_p}/{total_aeps} ({100*has_p//max(total_aeps,1)}%)')

# Missing CI (should have for key findings)
print('\nPapers with effect sizes but NO confidence intervals:')
for cid, aeps in sorted(all_aeps.items()):
    if aeps:
        no_ci = [a for a in aeps if a['evidence']['CI'] is None and
                 a['evidence']['metric'] in ('OR', 'AOR', 'aOR')]
        if no_ci:
            for a in no_ci[:2]:
                print(f'  [{cid}] {a["evidence"]["metric"]}={a["evidence"]["value"]} — no CI')

# ============================================================
# Save
# ============================================================
out_path = os.path.join(os.path.dirname(__file__), 'aep_database.json')
# Flatten for JSON
flat_aeps = []
for cid, aeps in all_aeps.items():
    for a in aeps:
        flat_aeps.append(a)

with open(out_path, 'w', encoding='utf-8') as f:
    json.dump({
        'total_aeps': total_aeps,
        'papers_with_aeps': sum(1 for v in all_aeps.values() if v),
        'report_issues': report_issues,
        'aeps': flat_aeps,
    }, f, ensure_ascii=False, indent=2)

print(f'\nSaved: {out_path} ({total_aeps} AEPs)')
print('Done.')
