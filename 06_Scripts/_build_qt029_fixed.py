# -*- coding: utf-8 -*-
"""Build FIXED QT029 tóm tắt docx — corrects HIGH-severity factual errors."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r"C:\Users\OS\OneDrive\read_books\Lo-au\Tom-tat-tung-bai\QT029_Li_CBT_NMA_2025_FIXED_07062026.docx"

doc = Document()

# Set base style: Times New Roman 12pt, line spacing 1.5
style = doc.styles['Normal']
style.font.name = 'Times New Roman'
style.font.size = Pt(12)
rpr = style.element.get_or_add_rPr()
rfonts = rpr.find(qn('w:rFonts'))
if rfonts is None:
    rfonts = OxmlElement('w:rFonts')
    rpr.append(rfonts)
rfonts.set(qn('w:ascii'), 'Times New Roman')
rfonts.set(qn('w:hAnsi'), 'Times New Roman')
rfonts.set(qn('w:cs'), 'Times New Roman')
style.paragraph_format.line_spacing = 1.5

def add_heading(text, bold=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p

def add_para(text):
    p = doc.add_paragraph(text)
    for r in p.runs:
        r.font.name = 'Times New Roman'
        r.font.size = Pt(12)
    p.paragraph_format.line_spacing = 1.5
    return p

# === TITLE ===
add_heading("TÓM TẮT BÀI QT029 (BẢN SỬA LỖI — 07/06/2026)")

# === Tên công trình ===
add_heading("Tên công trình, tác giả, năm")
add_para(
    'Công trình « Hiệu quả của các can thiệp khác nhau đối với rối loạn lo âu ở trẻ em '
    'và thanh thiếu niên: tổng quan hệ thống và phân tích tổng hợp mạng Bayes » '
    '(Effects of different interventions on anxiety disorders in children and adolescents: '
    'a systematic review and bayesian network meta-analysis).'
)
add_para(
    'Tác giả: Longhui Li, Qiner Li, Jingyi Wang, Quan Fu (đồng tác giả liên hệ, Capital '
    'University of Physical Education and Sports, Bắc Kinh), Meng Chi (đồng tác giả liên '
    'hệ, Shenyang Pharmaceutical University, Thẩm Dương, Trung Quốc).'
)
add_para(
    'Xuất bản trên BMC Psychiatry, 2025; 25:809. DOI: 10.1186/s12888-025-07227-y. '
    'Đăng ký PROSPERO: CRD42024587910. Truy cập mở (Open Access, CC BY-NC-ND 4.0). '
    '14 trang, 1 bảng (Table 1 — đặc điểm 30 RCTs), 6 hình (PRISMA, đánh giá chất '
    'lượng, sơ đồ mạng, biểu đồ hiệu quả tương đối, SUCRA, funnel plot).'
)

# === Mẫu khảo sát ===
add_heading("Mẫu khảo sát, đối tượng, địa bàn")
add_para(
    '30 thử nghiệm ngẫu nhiên có đối chứng (RCT) với 1.711 trẻ em và thanh thiếu niên '
    'từ 6 đến 18 tuổi, được chẩn đoán rối loạn lo âu theo DSM-5 (gồm ám ảnh sợ chuyên '
    'biệt, sợ khoảng rộng, rối loạn lo âu xã hội, rối loạn lo âu lan toả, rối loạn '
    'hoảng sợ và rối loạn lo âu chia ly). Tuổi trung bình 12,40 tuổi (qua 27 nghiên '
    'cứu báo cáo); 62% nữ (qua 29 nghiên cứu báo cáo). 30 RCTs đến từ 12 quốc gia: '
    'Mỹ (n=7), Úc (n=5), Đức (n=3), Anh (n=2), Iran (n=2), Thụy Điển (n=2) và các '
    'khu vực khác (n=9). Tuyển mộ từ bệnh viện (n=13), trường học (n=13) và cộng '
    'đồng (n=4). Không có RCT nào từ Việt Nam hay ASEAN.'
)

# === Phương pháp ===
add_heading("Phương pháp nghiên cứu")
add_para(
    'Tổng quan hệ thống và phân tích tổng hợp mạng Bayes (Bayesian network '
    'meta-analysis) tuân thủ hướng dẫn PRISMA, đăng ký trước trên PROSPERO. Phương '
    'pháp này cho phép so sánh đồng thời nhiều loại can thiệp ngay cả khi chưa từng '
    'được so sánh trực tiếp trong một RCT riêng lẻ, dựa trên mạng lưới bằng chứng '
    'trực tiếp và gián tiếp.'
)
add_para(
    'Tìm kiếm trong 5 cơ sở dữ liệu: Cochrane Library, Embase, PubMed, Scopus và '
    'Web of Science, giới hạn các bài công bố từ 01/01/1976 đến 01/09/2024. Bổ '
    'sung tìm tay qua Google Scholar và danh mục tham khảo của các tổng quan liên '
    'quan. Chỉ chọn bài tiếng Anh, có bình duyệt; loại tài liệu xám, luận án, kỷ '
    'yếu hội nghị.'
)
add_para(
    'Tiêu chí PICOS: (1) trẻ em/thanh thiếu niên 6–18 tuổi có chẩn đoán rối loạn '
    'lo âu DSM-5; loại trừ trẻ có bệnh đồng diễn nặng (rối loạn phổ tự kỷ, ADHD, '
    'ung thư, đái tháo đường); (2) can thiệp: trị liệu nhận thức–hành vi (CBT), '
    'trị liệu chấp nhận và cam kết (ACT), vận động thể chất (PE) hoặc trị liệu '
    'tiếp xúc thực tế ảo (VRET) — gồm cả các nghiên cứu kết hợp CBT với vận động; '
    '(3) đối chứng: chăm sóc thông thường, không can thiệp, hoặc danh sách chờ; '
    '(4) đầu ra chính: triệu chứng lo âu; (5) thiết kế: RCT.'
)
add_para(
    'Sàng lọc và trích xuất do hai nhà nghiên cứu độc lập (LLH, WJY), bất đồng do '
    'người thứ ba (LQE) giải quyết. Đánh giá nguy cơ thiên lệch bằng Cochrane '
    'Review Manager 5.4 theo Cochrane Handbook 5.1.0, với 6 chỉ số (phân nhóm '
    'ngẫu nhiên, che giấu phân nhóm, mù đôi, đầy đủ kết quả, báo cáo chọn lọc, '
    'thiên lệch khác). Phân tích NMA Bayes bằng gói Multinma 0.8.1 trong RStudio '
    '2024.12, với 4 chuỗi MCMC, mỗi chuỗi 2.000 vòng lặp (1.000 burn-in + 1.000 '
    'mẫu giữ lại). Cỡ hiệu lực biểu thị bằng chênh lệch trung bình (MD) và '
    'khoảng tin cậy 95% (95% CrI). Xếp hạng can thiệp bằng SUCRA. Đánh giá chất '
    'lượng bằng chứng theo GRADE (điều chỉnh cho NMA).'
)

# === Kết quả ===
add_heading("Phát hiện chính")
add_para(
    'Sàng lọc 19.442 bản ghi (cộng 6 bản ghi tìm tay), sau khi loại trùng còn '
    '10.596 bản, sàng lọc đầy đủ văn bản 79 bài, cuối cùng đưa vào 30 RCT. Hệ '
    'số kappa giữa hai người sàng lọc 0,83 (rất tốt).'
)
add_para(
    'Phân bố can thiệp trong 30 RCTs: CBT (n=16 nghiên cứu), ACT (n=6), VRET '
    '(n=6), PE (n=2). Thời lượng trung vị 10 tuần (3–24 tuần), 1 buổi/tuần '
    '(1–5), mỗi buổi 60 phút (20–120). Theo dõi trung vị 3 tháng (1–12).'
)
add_para(
    'Xếp hạng hiệu quả (so với nhóm đối chứng, MD và 95% CrI — Trang 7–8 PDF):'
)
add_para(
    '• Hạng 1 — ACT (Acceptance and Commitment Therapy): MD = −3,83 [95% CrI: '
    '−9,33; 1,51]; mean rank 2,25 [1; 5]; SUCRA = 0,69. Đây là can thiệp hiệu '
    'quả nhất.'
)
add_para(
    '• Hạng 2 — CBT (Cognitive Behavioral Therapy): MD = −3,64 [95% CrI: '
    '−7,36; −0,48]; mean rank 2,31 [1; 4]; SUCRA = 0,66. Đây là can thiệp duy '
    'nhất có khoảng tin cậy không bắt qua 0, cho thấy bằng chứng vững nhất về '
    'hiệu quả giảm lo âu.'
)
add_para(
    '• Hạng 3 — VRET (Virtual Reality Exposure Therapy): MD = −2,53 [95% CrI: '
    '−8,23; 3,32]; mean rank 2,86 [1; 5]; SUCRA = 0,51.'
)
add_para(
    '• Hạng 4 — PE (Physical Exercise): MD = −2,16 [95% CrI: −9,99; 5,52]; '
    'mean rank 3,12 [1; 5]; SUCRA = 0,51.'
)
add_para(
    'Tính nhất quán mạng lưới: node-splitting cho tất cả P > 0,05 (không có '
    'mâu thuẫn cục bộ). Mô hình ngẫu nhiên (DIC = 141,3) phù hợp hơn mô hình '
    'cố định (DIC = 263,4). Tính dị biệt đáng kể (τ = 5,87; 95% CrI: 2,92–9,83). '
    'Egger test p = 0,91 — không có thiên lệch xuất bản. So sánh head-to-head '
    'phổ biến nhất là CBT vs đối chứng (n = 16).'
)
add_para(
    'Đánh giá GRADE: chất lượng bằng chứng tổng thể THẤP do mâu thuẫn về tính '
    'không chính xác (imprecision — nhiều 95% CrI rất rộng và bắt qua 0) và '
    'tính dị biệt cao giữa các nghiên cứu.'
)
add_para(
    'Kết luận của tác giả: ACT có khả năng là can thiệp hiệu quả nhất cho rối '
    'loạn lo âu ở trẻ em/thanh thiếu niên, nhưng phát hiện cần được diễn giải '
    'thận trọng do chất lượng bằng chứng thấp, tính dị biệt cao và độ không '
    'chính xác lớn. Hướng nghiên cứu tương lai: kết hợp PE hoặc VR với trị liệu '
    'tâm lý; thêm RCT chất lượng cao quy mô lớn.'
)

# === Phản biện ===
add_heading("Phản biện")
add_para(
    '(1) Đính chính lỗi nghiêm trọng trong tóm tắt cũ: bài KHÔNG xếp hạng '
    '“CBT cá nhân hạng 1”. Bài so sánh 4 LOẠI can thiệp (ACT, CBT, VRET, PE), '
    'không phân tách CBT theo định dạng (cá nhân/nhóm/iCBT). ACT mới là can '
    'thiệp xếp hạng 1, CBT xếp hạng 2.'
)
add_para(
    '(2) Cảnh báo về độ tin cậy của xếp hạng: chênh lệch SUCRA giữa ACT (0,69) '
    'và CBT (0,66) chỉ 0,03 — gần như tương đương. Quan trọng hơn, 95% CrI của '
    'ACT (−9,33; 1,51) BẮT QUA 0, trong khi 95% CrI của CBT (−7,36; −0,48) '
    'KHÔNG bắt qua 0. Về mặt thống kê, CBT là can thiệp duy nhất có bằng chứng '
    'rõ ràng vượt đối chứng. Tuyên bố “ACT hiệu quả nhất” dựa trên điểm ước '
    'lượng MD, không phải ý nghĩa thống kê.'
)
add_para(
    '(3) Mất cân bằng số lượng nghiên cứu: ACT chỉ có 6 RCT, PE chỉ 2 RCT, '
    'trong khi CBT có 16 RCT. Xếp hạng ACT > CBT dựa trên cơ sở bằng chứng '
    'mỏng hơn đáng kể.'
)
add_para(
    '(4) GRADE THẤP cho toàn bộ mạng lưới — đây là cảnh báo nghiêm trọng. Tính '
    'dị biệt τ = 5,87 rất cao (so với MD chỉ −2 đến −4), nghĩa là khác biệt '
    'giữa các RCT lớn hơn cả hiệu quả can thiệp.'
)
add_para(
    '(5) Khoảng trống áp dụng cho Việt Nam: 30 RCTs đến từ 12 quốc gia — không '
    'có Việt Nam, không có ASEAN. Trung Quốc cũng chỉ 1 RCT (Yen 2014). Bằng '
    'chứng chủ yếu từ phương Tây (Mỹ + Úc + châu Âu ≈ 21/30 RCT).'
)
add_para(
    '(6) Hạn chế phương pháp do tác giả tự nhận: chỉ chọn bài tiếng Anh (loại '
    'tài liệu xám); không đánh giá kết quả lâu dài; không tách phân nhóm theo '
    'từng loại rối loạn lo âu (SAD vs GAD vs phobia).'
)

# === Áp dụng cho VN ===
add_heading("Áp dụng cho bối cảnh Việt Nam")
add_para(
    'Mặc dù tóm tắt cũ tuyên bố sai về xếp hạng, hàm ý chính sách vẫn có giá '
    'trị: CBT là can thiệp duy nhất có bằng chứng thống kê vững (95% CrI '
    'không bắt qua 0) và có cơ sở bằng chứng dày nhất (16 RCT). ACT là lựa '
    'chọn tiềm năng nhưng cần thêm bằng chứng. Không RCT nào từ Việt Nam — '
    'cần ưu tiên RCT CBT trong bối cảnh trường học hoặc bệnh viện nhi tại VN, '
    'có thể kết hợp PE (chi phí thấp, dễ triển khai) như tác giả gợi ý.'
)

# === Tham khảo ===
add_heading("Tham khảo")
add_para(
    'Li L, Li Q, Wang J, Fu Q, Chi M. Effects of different interventions on '
    'anxiety disorders in children and adolescents: a systematic review and '
    'bayesian network meta-analysis. BMC Psychiatry. 2025;25:809. '
    'DOI: 10.1186/s12888-025-07227-y. PROSPERO CRD42024587910. '
    'URL: https://doi.org/10.1186/s12888-025-07227-y'
)
add_para(
    'Truy vết nội bộ: PDF gốc tại 02_Papers-goc/The-gioi_Khac/'
    'QT029_BMC_Li_CBT_NMA_2025.pdf. Đối chiếu trực tiếp trang 1 (abstract, '
    'tác giả, DOI), trang 2–3 (intro, search strategy, PICOS), trang 4 '
    '(Cochrane RM 5.4, GRADE, MCMC), trang 6 (Table 1, n RCTs), trang 7–8 '
    '(SUCRA, MD, CrI, hạng từng can thiệp), trang 9 (figure 5 — SUCRA plot).'
)

# === Footer ===
add_heading("───────────────────────────────")
add_para(
    "Đã sửa lỗi factual (07/06/2026) — đối chiếu trực tiếp PDF gốc."
)

# === Strip metadata ===
cp = doc.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.category = ''
cp.last_modified_by = ''
cp.revision = 1

doc.save(OUT)
print("Wrote:", OUT)
print("Size:", os.path.getsize(OUT), "bytes")
