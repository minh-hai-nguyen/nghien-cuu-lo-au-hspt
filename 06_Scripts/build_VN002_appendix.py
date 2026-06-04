# -*- coding: utf-8 -*-
"""Part 5: Appendices 1-4 + References + 7-section Critique."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Cm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, '03_Ban-dich', 'VN002_VNAMHS_2022_National_FULL.docx')
IMG_DIR = 'C:/Users/HLC/AppData/Local/Temp/vnamhs_imgs/'

doc = Document(OUT)

def P(text='', bold=False, italic=False, size=None, color=None, align=None, red=False):
    p = doc.add_paragraph()
    if align is not None: p.alignment = align
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    if size: r.font.size = Pt(size)
    if bold: r.bold = True
    if italic: r.italic = True
    if red: r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)
    elif color: r.font.color.rgb = color
    return p

def H2(t): return P(t, bold=True, size=14, color=RGBColor(0x1F, 0x3A, 0x68))
def H3(t): return P(t, bold=True, size=12, color=RGBColor(0x2E, 0x54, 0x8B))
def H4(t): return P(t, bold=True, italic=True, size=11)

def PageMarker(page_num):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(f'--- Trang {page_num}, V-NAMHS, UNICEF/IOS 2022 ---')
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0xFF, 0x66, 0x00)

def AddImg(filename, caption_vn, caption_en, width_cm=11.0):
    path = os.path.join(IMG_DIR, filename)
    if not os.path.exists(path): return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    try: p.add_run().add_picture(path, width=Cm(width_cm))
    except: return
    cp = doc.add_paragraph()
    cp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cp.add_run(f'{caption_vn}\n({caption_en})')
    r.font.name = 'Times New Roman'; r.font.size = Pt(9); r.italic = True
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

def set_cell_bg(cell, color_hex):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd'); shd.set(qn('w:fill'), color_hex); tc_pr.append(shd)

def MakeTable(headers, rows, header_bg='D9E1F2'):
    t = doc.add_table(rows=1 + len(rows), cols=len(headers))
    t.style = 'Table Grid'
    for i, h in enumerate(headers):
        c = t.rows[0].cells[i]; c.text = ''
        r = c.paragraphs[0].add_run(h)
        r.font.name = 'Times New Roman'; r.font.size = Pt(10); r.bold = True
        set_cell_bg(c, header_bg)
    for ri, row in enumerate(rows):
        for ci, val in enumerate(row):
            c = t.rows[ri+1].cells[ci]; c.text = ''
            r = c.paragraphs[0].add_run(str(val))
            r.font.name = 'Times New Roman'; r.font.size = Pt(10)
    return t

# =============================================================
# APPENDIX 1: MEASURES
# =============================================================
PageMarker('36')
H2('PHỤ LỤC 1: Các công cụ đo lường (Appendix 1: Measures)')

P('Các công cụ đo được hỏi phụ huynh và vị thành niên được hiển thị dưới đây. Thông tin này được điều chỉnh từ Erskine và cộng sự (2021).')

H3('Phụ huynh (Parent)')
MakeTable(
    ['Công cụ', 'Mô tả'],
    [
        ('Demographics (Thông tin nhân khẩu)',
         'Thu thập thông tin nhân khẩu về hộ gia đình, vị thành niên và phụ huynh. Tiêu chí đủ điều kiện cũng được đánh giá trong module này theo các tiêu chí loại trừ nêu ở mục "Ai tham gia V-NAMHS?"'),
        ('Chronic illness (Bệnh mạn tính)',
         'Đo bệnh nghiêm trọng hoặc mạn tính mà vị thành niên hoặc phụ huynh/người chăm sóc đang trải qua.'),
        ('Pediatric Symptom Checklist – 17 (PSC-17)',
         'Bộ câu hỏi sàng lọc ngắn đánh giá triệu chứng nội hoá và ngoại hoá ở vị thành niên, được dùng để đo quan điểm của phụ huynh về SKTT của vị thành niên.'),
        ('Patient Health Questionnaire – 9 (PHQ-9)',
         'Công cụ sàng lọc ngắn để sàng lọc triệu chứng trầm cảm của phụ huynh.'),
        ('Generalised Anxiety Disorder – 7 (GAD-7)',
         'Công cụ sàng lọc ngắn để sàng lọc triệu chứng lo âu của phụ huynh.'),
        ('DISC-5: Introductory module',
         'Thiết lập dòng thời gian các sự kiện đáng kể trong 12 tháng qua để hỗ trợ người tham gia nhớ lại, và hướng dẫn cách trả lời các câu hỏi trong các module DISC-5.'),
        ('DISC-5: ADHD',
         'Đo tỷ lệ ADHD trong 12 tháng qua.'),
        ('Service use (Sử dụng dịch vụ)',
         'Thu thập thông tin từ phụ huynh về sử dụng dịch vụ, rào cản chăm sóc và nhu cầu được cảm nhận liên quan đến vị thành niên.'),
        ('COVID-19',
         'Đo tiếp xúc trực tiếp với COVID-19, kỳ thị, tác động kinh tế lên hộ gia đình, sử dụng chất của phụ huynh, và sử dụng dịch vụ của vị thành niên trong đại dịch COVID-19.'),
    ])

PageMarker('37')
H3('Vị thành niên (Adolescent)')
MakeTable(
    ['Công cụ', 'Mô tả'],
    [
        ('DISC-5: Introductory module',
         'Thiết lập dòng thời gian các sự kiện đáng kể trong 12 tháng qua để hỗ trợ người tham gia nhớ lại, và hướng dẫn cách trả lời các câu hỏi trong các module DISC-5.'),
        ('DISC-5: Social phobia (Ám ảnh xã hội)',
         'Đo tỷ lệ ám ảnh xã hội trong 12 tháng qua.'),
        ('DISC-5: Generalised anxiety disorder (Lo âu lan toả)',
         'Đo tỷ lệ rối loạn lo âu lan toả trong 12 tháng qua.'),
        ('DISC-5: Major depressive disorder (Trầm cảm nặng)',
         'Đo tỷ lệ rối loạn trầm cảm nặng trong 12 tháng qua. Bao gồm các câu hỏi về hành vi tự sát, được hỏi tất cả vị thành niên.'),
        ('Self-harm (Tự làm hại bản thân)',
         'Đo tỷ lệ, tuổi khởi phát, và gần đây nhất của tự làm hại bản thân.'),
        ('DISC-5: Conduct disorder (Rối loạn hành vi)',
         'Đo tỷ lệ rối loạn hành vi trong 12 tháng qua.'),
        ('DISC-5: PTSD',
         'Đo tỷ lệ PTSD trong 12 tháng qua.'),
        ('Informal help and self-help strategies',
         'Thu thập thông tin về hỗ trợ phi chính thức và chiến lược tự giúp.'),
        ('Self-rated health and body image',
         'Đo sức khoẻ tự đánh giá và hình ảnh cơ thể của vị thành niên.'),
        ('Physical activity',
         'Đo hoạt động thể chất của vị thành niên.'),
        ('Rosenberg Self-Esteem Scale',
         'Thang đo chuẩn ngắn về tự trọng (self-esteem).'),
        ('Bullying',
         'Đo tần suất bị bắt nạt và bắt nạt người khác, gồm cả phương thức bắt nạt.'),
        ('School and education',
         'Đo khát vọng học vấn (hiện tại và trong quá khứ tuỳ vào tình trạng học hiện tại), kỳ vọng và áp lực.'),
        ('Peer relationships and loneliness',
         'Thu thập thông tin về tình bạn của vị thành niên (gồm peer deviance) và sự cô đơn.'),
        ('GEAS Family Connectedness',
         'Thu thập thông tin về quan hệ của vị thành niên với phụ huynh.'),
        ('Religiosity',
         'Đo hỗ trợ được cảm nhận từ cộng đồng tín ngưỡng.'),
        ('Safety and security',
         'Đo sự an toàn cá nhân được cảm nhận trong các bối cảnh khác nhau ví dụ nhà, trường, hàng xóm.'),
        ('Sexual behaviour *',
         'Thu thập thông tin về hành vi tình dục của vị thành niên, tính dục, và bản dạng giới. Chỉ hỏi VTN 12–17 tuổi.'),
        ('Adverse Childhood Experiences (ACEs) questionnaire *',
         'Đo mức độ tiếp xúc suốt đời với nhiều loại lạm dụng, bỏ rơi, bạo lực giữa bố mẹ/người chăm sóc, các loại rối loạn chức năng hộ gia đình nghiêm trọng khác, và bạo lực.'),
        ('Substance use *',
         'Đo sử dụng thuốc lá, rượu, cần sa và các chất ma tuý bất hợp pháp khác.'),
        ('COVID-19',
         'Đo tiếp xúc trực tiếp với COVID-19, tác động giáo dục, nghịch cảnh hộ gia đình/cá nhân, và vấn đề cảm xúc và hành vi trong đại dịch COVID-19.'),
    ])
P('* Các module này được vị thành niên tự trả lời (self-administered).',
  italic=True, size=9, color=RGBColor(0x80, 0x80, 0x80))

# =============================================================
# APPENDIX 2: METHODOLOGY
# =============================================================
PageMarker('38')
H2('PHỤ LỤC 2: Phương pháp (Appendix 2: Methodology)')

H3('Khung mẫu (Sampling frame)')
P('Mục tiêu chính của khung mẫu V-NAMHS là cung cấp ước tính thống kê đáng tin cậy về tỷ lệ rối loạn tâm thần ở vị thành niên. Mẫu mang tính đại diện quốc gia, bao gồm khu vực đô thị và nông thôn.')

P('Khung mẫu V-NAMHS bao gồm cả 63 tỉnh/thành phố, gồm 713 huyện, được chia thành 4 vùng (Trung du miền núi phía Bắc và Tây Nguyên; Đồng bằng sông Hồng; Bắc Trung Bộ và Duyên hải miền Trung; Đông Nam Bộ và Đồng bằng sông Cửu Long).')

MakeTable(
    ['Vùng', 'Dân số (người)', '% tổng DS', 'Đô thị (người)', '% đô thị của vùng'],
    [
        ('Vùng 1: Trung du miền núi phía Bắc + Tây Nguyên', '18.375.547', '19,1', '3.957.095', '21,5'),
        ('Vùng 2: Đồng bằng sông Hồng', '22.543.607', '23,4', '7.856.566', '34,9'),
        ('Vùng 3: Bắc Trung Bộ + Duyên hải miền Trung', '20.187.293', '21,0', '5.719.511', '28,3'),
        ('Vùng 4: Đông Nam Bộ + ĐBSCL', '35.102.537', '36,5', '15.526.563', '44,2'),
        ('Tổng', '96.208.984', '100,0', '33.059.735', '34,4'),
    ])
P('Bảng 18. Dân số Việt Nam năm 2019 theo vùng và khu vực đô thị – nông thôn. Nguồn: Tổng cục Thống kê 2020.',
  italic=True, size=9, color=RGBColor(0x66, 0x66, 0x66))

P('Tổng cộng, 38 tỉnh/thành phố từ 4 vùng được chọn tham gia khảo sát (xem bản đồ ở Hình 5), gồm:')
P('• Vùng 1: Hà Giang, Cao Bằng, Bắc Kạn, Tuyên Quang, Điện Biên, Yên Bái, Thái Nguyên, Bắc Giang, Kon Tum, Đắk Lắk, Lâm Đồng.')
P('• Vùng 2: Hà Nội, Quảng Ninh, Vĩnh Phúc, Hải Dương, Hải Phòng, Hưng Yên, Thái Bình, Nam Định.')
P('• Vùng 3: Thanh Hoá, Nghệ An, Quảng Bình, Đà Nẵng, Quảng Ngãi, Phú Yên, Ninh Thuận.')
P('• Vùng 4: Bình Phước, Bình Dương, Đồng Nai, Bà Rịa – Vũng Tàu, TP Hồ Chí Minh, Tiền Giang, Trà Vinh, Đồng Tháp, An Giang, Kiên Giang, Hậu Giang, Bạc Liêu.')

P('Hộ gia đình tại 38 tỉnh được chọn được nhóm thành các đơn vị điều tra gọi là Địa bàn Điều tra (Enumeration Area — EA), mỗi EA có trung bình khoảng 120 hộ. Sau đó 200 EA được chọn ngẫu nhiên từ 38 tỉnh, phân đều qua 4 vùng — tức 50 EA mỗi vùng (25 đô thị + 25 nông thôn). Số hộ được chọn cho mỗi EA trong V-NAMHS là 38 hộ. Tỷ lệ trả lời dự kiến 79 % (tỷ lệ không trả lời = 21 %) được áp dụng khi tính cỡ mẫu V-NAMHS. Việc chọn 7.600 hộ tương ứng với con số kỳ vọng 6.000 hộ hoàn thành phỏng vấn.')

PageMarker('39')
AddImg('p041_img1_Image257.jpg',
       'Hình 5. Bản đồ Việt Nam và 38 tỉnh/thành phố được chọn',
       'Figure 5. Map of Viet Nam and 38 selected provinces/cities',
       width_cm=12.0)

H3('Nghiên cứu thí điểm (Pilot study)')
P('Trước khi bắt đầu thu thập dữ liệu V-NAMHS, một nghiên cứu thí điểm được tiến hành để thử nghiệm tất cả quy trình và logistics. Thí điểm được thực hiện tại Hà Nội (miền Bắc) và Đồng Nai (miền Nam). Tại Hà Nội, 25 vị thành niên (9 nam và 16 nữ) cùng phụ huynh đã được phỏng vấn. Tại Đồng Nai, thêm 25 vị thành niên (12 nam và 13 nữ) cùng phụ huynh được phỏng vấn. Thí điểm nhằm kiểm tra tính nhất quán và mạch lạc của câu hỏi nghiên cứu, độ dài phỏng vấn, phù hợp ngôn ngữ, và cách đặt các câu hỏi khó hoặc nhạy cảm. Thí điểm cũng dùng để kiểm tra công cụ được lập trình trên nền tảng SurveyCTO và thông tin cho kế hoạch và tổ chức khảo sát. Hầu hết thách thức liên quan đến độ dài công cụ, khó khăn hiểu của VTN nhỏ hơn, và các vấn đề lập trình. Các thách thức này được rà soát và thực hiện các chỉnh sửa cần thiết trước khi thu thập dữ liệu.')

PageMarker('40')
H3('Thực địa (Fieldwork)')
P('Thu thập dữ liệu tại VN bắt đầu từ ngày 21/9/2022 [lưu ý: bản gốc ghi "2022" ở Phụ lục 2 nhưng ở phần Introduction ghi "2021" — có khả năng lỗi in ấn trong bản gốc; ngày hoàn tất 16/12/2021 cho thấy đúng là năm 2021], bắt đầu miền Bắc rồi miền Nam. Thu thập dữ liệu có tổng 127 phỏng vấn viên tham gia. Công tác thực địa được triển khai theo quy trình thực địa (fieldwork protocol) được phát triển và thiết lập bởi cả năm nhóm NAMHS quốc tế, rồi áp dụng tại VN. Nhóm V-NAMHS phối hợp chặt chẽ với GOPFP để xác định ranh giới trong quá trình liệt kê và lập bản đồ hộ gia đình ở 38 tỉnh. Tổng 200 EA được chọn, mỗi EA chứa khoảng 120 hộ trong đó 38 hộ được chọn cho mẫu V-NAMHS. Do hạn chế di chuyển liên quan COVID-19 giữa các tỉnh, phỏng vấn viên là nhân viên của POPFP địa phương, tránh việc di chuyển giữa các tỉnh. Số phỏng vấn viên cho mỗi tỉnh phụ thuộc kích thước tỉnh, dao động từ 2 đến 16 người. Thu thập dữ liệu chỉ được tiến hành nếu EA được chọn được phân loại "xanh" (an toàn) hoặc "vàng" (nguy cơ COVID-19 thấp), tạm ngừng nếu EA là "cam" (nguy cơ cao) hoặc "đỏ" (nguy cơ rất cao).')

P('Trong các tuần trước thu thập dữ liệu, thư thông tin về V-NAMHS được cung cấp cho các hộ được chọn trong mỗi EA. Các quy trình đảm bảo tham gia tự nguyện của phụ huynh và vị thành niên, quyền riêng tư, an toàn (cho phỏng vấn viên và người tham gia) và an toàn dữ liệu cũng được bao gồm trong thông tin cung cấp cho hộ. Điều này được hỗ trợ bởi sự hợp tác chặt chẽ giữa nhóm V-NAMHS, GOPFP và các POPFP.')

P('Trong suốt thu thập dữ liệu, giám sát và hỗ trợ kỹ thuật được nhóm V-NAMHS cung cấp cho nhân viên thực địa hàng ngày qua liên lạc trực tuyến. Nhóm V-NAMHS cũng phối hợp với UQ và JHSPH để giám sát dữ liệu theo quy trình định trước. Tư vấn kỹ thuật và xử lý sự cố cũng được các nhóm này cung cấp liên tục.')

# =============================================================
# APPENDIX 3: GLOSSARY
# =============================================================
PageMarker('41')
H2('PHỤ LỤC 3: Thuật ngữ (Appendix 3: Glossary)')

MakeTable(
    ['Thuật ngữ', 'Định nghĩa'],
    [
        ('Tỷ lệ 12 tháng qua (12-month prevalence)',
         'Đáp ứng tiêu chí của một vấn đề SKTT hoặc rối loạn tâm thần theo DISC-5 trong 12 tháng trước khi phỏng vấn. Bao gồm những người có triệu chứng khởi phát trong 12 tháng trước khi phỏng vấn, và những người có triệu chứng khởi phát sớm hơn nhưng tiếp tục đáp ứng tiêu chí trong 12 tháng qua.'),
        ('Vị thành niên (Adolescent)',
         'Người trẻ 10–17 tuổi. Dù WHO định nghĩa vị thành niên là 10–19 tuổi, các em 18–19 tuổi bị loại khỏi nghiên cứu vì nhiều khả năng đã sống độc lập và/hoặc làm việc xa nhà. Thêm nữa, công cụ chẩn đoán (như DISC-5) không được thiết kế cho người ≥ 18 tuổi — thường được đánh giá bằng công cụ dành cho người lớn.'),
        ('Các rối loạn lo âu (Anxiety disorders)',
         'Một nhóm rối loạn tâm thần định nghĩa bởi nỗi sợ và lo âu quá mức. Ám ảnh xã hội và rối loạn lo âu lan toả là hai rối loạn lo âu được bao gồm trong khảo sát này.'),
        ('ADHD (Attention-deficit/hyperactivity disorder)',
         'Đặc trưng bởi các mô hình dai dẳng của mất chú ý và/hoặc tăng động – bốc đồng. Vị thành niên có thể gặp khó khăn về chú ý và tập trung, có cử động quá mức và/hoặc khó kiểm soát hành vi bốc đồng. Các hành vi này không tương xứng với độ tuổi hoặc mức độ phát triển của VTN và xảy ra ở nhiều bối cảnh. Module ADHD của DISC-5 được hỏi phụ huynh.'),
        ('Rối loạn hành vi (Conduct disorder)',
         'Đặc trưng bởi mô hình hành vi lặp đi lặp lại vi phạm quyền của người khác và/hoặc các quy tắc hoặc chuẩn mực xã hội lớn. Hành vi có thể bao gồm hung hăng với người hoặc động vật, phá hoại tài sản, lừa dối hoặc trộm cắp, hoặc vi phạm nghiêm trọng quy tắc. Module rối loạn hành vi của DISC-5 được hỏi vị thành niên.'),
        ('Tiêu chí chẩn đoán (Diagnostic criteria)',
         'Tập hợp yêu cầu cụ thể mà vị thành niên phải đáp ứng để được xem là có rối loạn tâm thần. Tiêu chí có thể bao gồm: tập hợp số hoặc tổ hợp triệu chứng cụ thể; tuổi khởi phát triệu chứng hoặc hành vi; tần suất và thời gian của triệu chứng; căng thẳng hoặc suy giảm chức năng. Trong V-NAMHS, tiêu chí chẩn đoán được xác định theo DSM-5.'),
        ('DISC-5 (Diagnostic Interview Schedule for Children, Version 5)',
         'Công cụ chẩn đoán có cấu trúc hoàn toàn được thiết kế để xác định trẻ em hoặc vị thành niên đáp ứng tiêu chí chẩn đoán DSM-5. Sáu module chẩn đoán từ DISC-5 được bao gồm trong khảo sát và điều chỉnh để đảm bảo phù hợp văn hoá đồng thời duy trì tính nhất quán về khái niệm.'),
        ('DSM-5 (Diagnostic and Statistical Manual of Mental Disorders, 5th Edition)',
         'Định nghĩa của từng rối loạn tâm thần được Hiệp hội Tâm thần Hoa Kỳ xuất bản và dùng để định nghĩa và chẩn đoán rối loạn tâm thần.'),
        ('Triệu chứng đầy đủ ngưỡng (Full threshold symptoms)',
         'Ghi nhận tất cả các triệu chứng cần thiết để đáp ứng tiêu chí chẩn đoán DSM-5 cho một rối loạn tâm thần (lưu ý suy giảm chức năng cũng phải được ghi nhận để đáp ứng tiêu chí DSM-5).'),
        ('Rối loạn lo âu lan toả (Generalised anxiety disorder)',
         'Đặc trưng bởi lo âu và lo lắng quá mức về nhiều sự kiện hoặc hoạt động. Cường độ, tần suất và thời gian lo âu không tương xứng với khả năng thực tế hoặc tác động của sự kiện dự kiến. Module này của DISC-5 được hỏi vị thành niên.'),
        ('Suy giảm chức năng (Impairment)',
         'Khi triệu chứng của một rối loạn tâm thần ảnh hưởng tiêu cực hoặc can thiệp vào hoạt động và/hoặc các khía cạnh khác nhau của cuộc sống vị thành niên. Ghi nhận suy giảm là cần thiết để đáp ứng tiêu chí chẩn đoán DSM-5 cho một rối loạn tâm thần (cùng với ghi nhận tất cả triệu chứng cần thiết — tức triệu chứng đầy đủ ngưỡng).'),
        ('Lĩnh vực suy giảm (Impairment domains)',
         'Trong DISC-5, suy giảm được đánh giá bằng 6 câu hỏi đo suy giảm do triệu chứng trên 4 lĩnh vực: gia đình, bạn bè, trường hoặc việc làm, căng thẳng cá nhân.'),
        ('Rối loạn trầm cảm nặng (Major depressive disorder)',
         'Đặc trưng bởi giai đoạn ít nhất 2 tuần trong đó có tâm trạng trầm, mất quan tâm hoặc vui vẻ trong gần như tất cả hoạt động, và/hoặc khó chịu. Những cảm giác này cũng gắn với các triệu chứng thể chất khác như mệt mỏi, rối loạn giấc ngủ hoặc vấn đề tập trung. Module này của DISC-5 được hỏi vị thành niên.'),
        ('Rối loạn tâm thần (Mental disorder)',
         'Là hội chứng hành vi hoặc tâm lý có ý nghĩa lâm sàng xảy ra ở một cá nhân và gắn với căng thẳng hiện tại (triệu chứng đau đớn), khuyết tật (suy giảm ở một hoặc nhiều lĩnh vực chức năng quan trọng), và/hoặc nguy cơ tử vong, đau đớn, khuyết tật hoặc mất tự do đáng kể tăng lên. Trong báo cáo này, VTN có rối loạn tâm thần là những em đáp ứng tiêu chí chẩn đoán DSM-5 cho một rối loạn tâm thần cụ thể đo trong V-NAMHS.'),
        ('Vấn đề SKTT (Mental health problem)',
         'Tương tự rối loạn tâm thần ở chỗ cũng can thiệp vào cách một người suy nghĩ, cảm xúc và hành vi, nhưng ở mức độ ít hơn so với rối loạn tâm thần. Có thể trải nghiệm tạm thời, hoặc như một phản ứng cấp tính với căng thẳng cuộc sống. Trong báo cáo này, VTN có vấn đề SKTT bao gồm các em đáp ứng chẩn đoán rối loạn DSM-5 (tức triệu chứng đầy đủ ngưỡng + ghi nhận suy giảm) cũng như các em không ghi nhận suy giảm (triệu chứng đầy đủ nhưng không suy giảm) và các em đáp ứng ít nhất một nửa số triệu chứng DSM-5 yêu cầu (tức subthreshold symptoms) có hoặc không suy giảm.'),
        ('Phụ huynh (Parent)',
         'Trong báo cáo này, dùng để chỉ người được đề cử là người chăm sóc chính của vị thành niên (primary caregiver).'),
        ('PTSD (Posttraumatic stress disorder)',
         'Đặc trưng bởi các ý nghĩ xâm nhập hoặc lặp lại, phân ly, nhận thức bóp méo hoặc tiêu cực, tăng kích thích hoặc phản ứng, hoặc các triệu chứng xâm nhập khác hoặc phản ứng thể chất, tất cả liên quan đến một chấn thương cụ thể. Module này của DISC-5 được hỏi vị thành niên.'),
        ('Người chăm sóc chính (Primary caregiver)',
         'Người có trách nhiệm, chăm sóc, và có khả năng cung cấp thông tin tốt nhất về vị thành niên. Người chăm sóc chính tự nhận mình vào đầu phỏng vấn. Trong báo cáo này được gọi là "phụ huynh".'),
        ('Tự làm hại bản thân (Self-harm)',
         'Hành động cố ý gây tổn hại hoặc chấn thương cho bản thân mà không có ý định kết thúc cuộc sống. Điều này phân biệt tự làm hại với toan tự sát.'),
        ('Dịch vụ (Service)',
         'Trong khảo sát này, dịch vụ được xem là bất kỳ nhà cung cấp nào cung cấp hỗ trợ hoặc tư vấn cho vấn đề cảm xúc và hành vi. Các nhà cung cấp dịch vụ được bao gồm trong khảo sát: bác sĩ hoặc điều dưỡng; chuyên gia (ví dụ nhà tâm lý học hoặc bác sĩ tâm thần); nhân viên y tế cộng đồng; nhân viên trường học (giáo viên, huấn luyện viên, tư vấn viên trường học); lãnh đạo tôn giáo/tín ngưỡng; thầy thuốc truyền thống; Khác (theo định nghĩa của người tham gia). Định nghĩa nhà cung cấp được mở rộng để bao gồm những người thường không được xem là nhà cung cấp do khả năng các ngành này được tiếp cận cho các dịch vụ như vậy.'),
        ('Sử dụng dịch vụ (Service use)',
         'Được định nghĩa là sử dụng bất kỳ dịch vụ nào (bởi các nhà cung cấp liệt kê trên) cho hỗ trợ hoặc tư vấn cho vấn đề cảm xúc và hành vi. Câu hỏi sử dụng dịch vụ được hỏi phụ huynh.'),
        ('Ám ảnh xã hội (Social phobia)',
         'Đặc trưng bởi nỗi sợ một hoặc nhiều tình huống xã hội, nơi vị thành niên là trung tâm của sự chú ý của người khác, có thể gây ra cảm giác xấu hổ và nhục nhã. Điều này có thể dẫn đến việc vị thành niên tránh né các tình huống hoặc chịu đựng chúng nhưng sợ hãi khi làm vậy. Ở vị thành niên, các tình huống gây lo âu phải ở môi trường peer, không chỉ xung quanh người lớn. Module này của DISC-5 được hỏi vị thành niên.'),
        ('Triệu chứng dưới ngưỡng (Subthreshold symptoms)',
         'Trong DISC-5, một vị thành niên được xem là có triệu chứng dưới ngưỡng nếu ghi nhận ít nhất một nửa số triệu chứng DSM-5 yêu cầu nhưng không phải tất cả.'),
        ('Hành vi tự sát (Suicidal behaviours)',
         'Bao gồm ý nghĩ tự sát, lập kế hoạch tự sát, và toan tự sát.'),
        ('Ý nghĩ tự sát (Suicidal ideation)',
         'Nghĩ về việc muốn chết hoặc ý nghĩ chung về kết thúc cuộc sống của mình.'),
        ('Toan tự sát (Suicide attempt)',
         'Tự gây hại với ý định kết thúc cuộc sống của mình.'),
        ('Lập kế hoạch tự sát (Suicide planning)',
         'Lập kế hoạch để kết thúc cuộc sống của mình.'),
    ])

# =============================================================
# APPENDIX 4: RESEARCH TEAMS
# =============================================================
PageMarker('44')
H2('PHỤ LỤC 4: Nhóm nghiên cứu (Appendix 4: Research teams)')
P('Tên các thành viên giữ nguyên tiếng Anh/Việt như trong bản gốc để bảo toàn thông tin và credit.',
  italic=True, size=10, color=RGBColor(0x80, 0x80, 0x80))

P('Viện Xã hội học (IOS) – Chủ trì tại VN', bold=True)
P('TS. Vũ Mạnh Lợi — V-NAMHS Principal Investigator (VN)')
P('TS. Nguyễn Đức Vinh — V-NAMHS Senior Research Officer')
P('TS. Đào Thị Khánh Hoa — V-NAMHS Senior Research Officer')
P('Prof. Đặng Nguyên Anh — V-NAMHS Senior Research Officer')
P('TS. Nghiêm Thị Thuỷ — V-NAMHS Administrator')
P('PGS.TS. Nguyễn Đức Chiến — V-NAMHS Research Officer')
P('TS. Hoàng Vũ Linh Chi — V-NAMHS Research Officer')
P('Khuất Thị Diệu Linh — V-NAMHS Research Officer')
P('Trần Việt Long — V-NAMHS Research Officer')
P('Nguyễn Thị Xuân — V-NAMHS Research Officer')
P('Nguyễn Quang Tuấn — V-NAMHS Research Officer')

P('Nhóm GOPFP', bold=True)
P('TS. Phạm Vũ Hoàng — GOPFP Team Leader')
P('Nguyễn Cao Trường — Fieldwork Administrator')

P('University of Queensland (UQ)', bold=True)
P('Dr. Holly Erskine — NAMHS Principal Investigator')
P('Prof. Harvey Whiteford — NAMHS Senior Advisor')
P('Prof. James Scott — NAMHS Clinical Advisor')
P('Dr. Sarah Blondell — NAMHS Senior Research Officer')
P('Krystina Wallis — NAMHS Research Officer')
P('Cartiah McGrath — NAMHS Research Officer')

P('Johns Hopkins Bloomberg School of Public Health (JHSPH)', bold=True)
P('Prof. Robert Blum — JHSPH NAMHS Project Lead')
P('Dr. Shoshanna Fine — JHSPH NAMHS Assistant Scientist')
P('Mengmeng Li — JHSPH NAMHS Data Analyst')
P('Astha Ramaiya — JHSPH NAMHS Assistant Scientist')

# image on p51
AddImg('p051_img1_Image304.jpg',
       'Hình bìa cuối — V-NAMHS 2022',
       'Back cover image — V-NAMHS 2022',
       width_cm=10.0)

doc.save(OUT)
print(f'Part 5 (Appendices 1-4) saved.')
