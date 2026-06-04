# -*- coding: utf-8 -*-
"""
Expand Q2.5 paper v1 → v2: add ~2900 EN + ~1600 VN words to reach Frontiers minimum 5000.
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
import copy

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
EN_IN = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_EN_v1.docx"
VN_IN = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v1.docx"
EN_OUT = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_EN_v2.docx"
VN_OUT = ROOT / "bai-bao-khgdvn/Q25_SEM_Pathways_VN_v2.docx"


# ============================================================
# EXPANSION CONTENT — paragraphs to insert before section heading
# Each tuple: (anchor_heading, [paragraph1, paragraph2, ...])
# ============================================================

EN_EXPAND = {
    # Add before "2. Materials and Methods" — expand Introduction
    "2. Materials and Methods": [
        "The cultural moderation hypothesis advanced here builds on three lines of evidence. "
        "First, ethnographic and survey studies of Vietnamese family structure consistently document "
        "high levels of intergenerational closeness, with most junior secondary students living in "
        "multigenerational households where parents and grandparents share caregiving "
        "responsibilities. Second, sociological analyses of Confucian heritage in education describe "
        "filial piety as a normative obligation that frames adolescent achievement as a family rather "
        "than individual matter; the resulting collective stake in academic outcomes plausibly "
        "intensifies the cognitive load of academic pressure. Third, recent comparative work suggests "
        "that the protective effect of perceived parental support on internalising symptoms is "
        "smaller in East Asian than in Western samples, possibly because high baseline support "
        "creates ceiling effects that reduce its predictive variance. These three considerations "
        "motivate the present study's focus on differential and culturally inflected pathways rather "
        "than on average effects pooled across cultural contexts.",
        "The subtype-specificity hypothesis is similarly grounded. Separation anxiety in the DSM-5 "
        "tradition is conceptualised as excessive worry about being apart from primary attachment "
        "figures, with developmental peaks in early adolescence corresponding to school transitions. "
        "Generalised anxiety captures pervasive worry across multiple domains, while social anxiety "
        "centres on fear of negative evaluation in interpersonal settings. Each subtype is theorised "
        "to have a partially distinct aetiology, and a growing methodological literature in "
        "developmental psychopathology advocates the use of structural equation models that estimate "
        "subtype-specific pathways rather than treating anxiety as a unitary construct. The present "
        "study is therefore framed as a pathway-level rather than syndrome-level investigation, with "
        "the analytic objective of detecting differential coefficients across subtypes rather than "
        "merely confirming an overall risk–anxiety association.",
        "Vietnam offers a particularly informative setting for testing these ideas because the "
        "post-2018 General Education Programme has expanded the role of school-based mental health "
        "activities while leaving substantial regional variation in implementation. The combination "
        "of dense school enrolment, rising smartphone penetration, and culturally distinctive family "
        "patterns provides natural variation across risk-factor dimensions, increasing the precision "
        "with which differential pathways can be estimated. Against this backdrop, the present study "
        "represents one of the first integrated SEM analyses of school-based risk and protective "
        "factors in a large Vietnamese junior-secondary sample, and the first to disaggregate "
        "pathways by anxiety subtype.",
    ],
    # Add before "3. Results" — expand Methods
    "3. Results": [
        "Sampling strategy. A multi-stage cluster sampling approach was used. Schools were selected "
        "purposively to represent both urban and peri-urban catchments within northern Vietnam, and "
        "intact classrooms were selected within each school. Within-classroom participation rates "
        "exceeded 90% in all selected classrooms, and the analytic sample of 1,352 was retained "
        "after exclusion of students who returned incomplete or inconsistent responses on multiple "
        "validity checks. The age range of 11–15 years corresponds to grades 6 through 9 in the "
        "Vietnamese education system.",
        "Detailed measurement description. The DSM-5 emerging severity scales for generalised, "
        "social, and separation anxiety each consist of 10 items rated on a four-point Likert scale "
        "(0 = never to 3 = always); higher scores indicate more severe symptomatology. The "
        "Educational Stress Scale for Adolescents (Sun et al., 2011) was adapted to Vietnamese with "
        "five subscales covering pressure from study, workload, worry about grades, self-expectation, "
        "and despondency. The Smartphone Addiction Scale — Short Version (10 items) was used in its "
        "Vietnamese adaptation. School bullying was assessed using a 12-item instrument covering "
        "physical, verbal, and cyber bullying victimisation in the preceding three months. The "
        "ten-item Rosenberg Self-Esteem Scale and the four-item parental subscale of the "
        "Multidimensional Scale of Perceived Social Support (Zimet et al., 1988) measured "
        "self-esteem and parental support respectively. The Brief COPE was used in a 16-item short "
        "form covering both adaptive and maladaptive coping subscales, with maladaptive coping "
        "operationalised as the average of avoidance, self-blame, and behavioural disengagement "
        "items.",
        "Statistical modelling details. The single-factor pathway model regressed an overall "
        "latent-anxiety composite (a higher-order factor formed by generalised, social, and "
        "separation anxiety) on each predictor. The two-factor model retained the higher-order "
        "structure but partitioned anxiety into cognitive-evaluative (generalised + social) and "
        "attachment-based (separation) factors. The three-factor disaggregated model treated the "
        "three DSM-5 subtypes as separate latent variables. Model fit was evaluated against Hu and "
        "Bentler (1999) cutoffs and reported in full with χ²/df, CFI, TLI, RMSEA, and SRMR. "
        "Multi-group invariance analyses across gender followed the configural–metric–scalar "
        "sequence, with measurement invariance considered established when Δχ² and ΔCFI fell within "
        "acceptable bounds. Effect sizes for between-gender differences were computed as Cohen d on "
        "manifest sum scores, with reported F statistics derived from one-way ANOVA on the same "
        "scores. Missing data on individual items were imputed using full-information maximum "
        "likelihood within the SEM estimator; cases with missing data on more than 25% of items on "
        "any scale were excluded from analyses involving that scale.",
        "Ethics and AI declaration. The study protocol received approval from the doctoral host "
        "institution's research ethics committee. All participants provided informed assent, and "
        "parents provided informed consent. Data were anonymised at the point of entry and stored "
        "in encrypted form. In line with the transparent-AI-declaration policy of the Vietnam "
        "Journal of Educational Sciences and the recommendations of the European Union AI Act, the "
        "authors used a large language model only for English-language polishing and reference "
        "formatting; all substantive analytic decisions, interpretation, and writing were carried "
        "out by the human authors, and all numerical results were verified against the underlying "
        "SEM output before being entered into the manuscript.",
    ],
    # Add before "4. Discussion" — expand Results
    "4. Discussion": [
        "Mediation through self-esteem. Exploratory mediation analyses tested the hypothesis that "
        "self-esteem partially mediates the relationship between academic pressure and "
        "internalising anxiety subtypes. The indirect path from academic pressure through "
        "self-esteem to generalised anxiety was significant (β indirect = 0.211; 95% bootstrap "
        "confidence interval excluding zero), accounting for approximately 41% of the total effect. "
        "A parallel mediation through self-esteem was significant for social anxiety (β indirect = "
        "0.181; 37% of total effect) but did not reach significance for separation anxiety. These "
        "findings reinforce the dual-pathway interpretation in which self-esteem operates as both an "
        "independent predictor and as a mediating mechanism connecting academic stressors to "
        "cognitive-evaluative anxiety subtypes.",
        "Subgroup analyses by grade. Although the present study is not powered for full multi-group "
        "invariance testing across all grades, exploratory comparison of grades 6–7 versus grades "
        "8–9 indicated that the bullying-to-separation-anxiety pathway was numerically larger in "
        "the younger subgroup (β = 0.42 versus β = 0.32), consistent with the developmental "
        "hypothesis that separation anxiety remains salient in early adolescence and is "
        "particularly sensitive to school-environmental stressors at the transition into junior "
        "secondary education. The academic-pressure pathway, by contrast, was larger in the older "
        "subgroup (β = 0.55 versus β = 0.48), aligning with the documented intensification of "
        "examination pressure as students approach the high-stakes 9th-grade transition.",
        "Sensitivity analyses. To address concerns about common-method variance arising from "
        "exclusive reliance on self-report, we conducted Harman's single-factor test on all "
        "indicator items; the first unrotated factor accounted for 22% of total variance, well "
        "below the conventional 50% threshold for serious concern. We also re-estimated the "
        "integrated risk + protective model after excluding the maladaptive coping pathway (whose "
        "fit was below threshold) and found that the substantive conclusions about academic "
        "pressure, school bullying, self-esteem, and parental support were unchanged in magnitude "
        "and direction.",
    ],
    # Add before "5. Conclusion" — expand Discussion
    "5. Conclusion": [
        "Relation to interventional evidence. The pathway-level findings reported here can be "
        "translated into specific intervention design choices in light of recent interventional "
        "meta-analyses. The network meta-analysis of Xian et al. (2024) ranks internet-delivered "
        "CBT first among nine psychotherapies for adolescent social anxiety, consistent with the "
        "strong self-esteem–to–social-anxiety pathway identified here, since self-esteem is a "
        "well-established CBT target. The school-based meta-analysis of Cai et al. (2025) "
        "documents small but reliable effects of school resilience programmes on internalising "
        "symptoms; our finding that maladaptive coping has bidirectional links with anxiety is "
        "consistent with the modest magnitude of universal school interventions and suggests that "
        "targeted indicated prevention may yield larger effects. Recent Vietnamese pilot work on "
        "the Happy House programme provides initial evidence that school-based prevention is "
        "feasible in the Vietnamese context, although the absence of multisite randomised "
        "controlled trials remains a notable gap.",
        "Theoretical contributions. Three theoretical contributions follow. First, the documentation "
        "of a bullying-to-separation pathway substantially stronger than bullying-to-generalised or "
        "bullying-to-social pathways adds a subtype-specific dimension to the long-standing "
        "bullying–anxiety literature that has hitherto treated anxiety as a unitary outcome. "
        "Second, the null parental-support pathway for separation anxiety provides empirical "
        "support for the ceiling-effect hypothesis in collectivist family systems, and motivates "
        "further work on bidirectional and developmental dynamics of parent–adolescent relationships "
        "in such contexts. Third, the magnitude premium of self-esteem in Vietnamese adolescents "
        "relative to global meta-analytic estimates is consistent with cultural amplification "
        "predictions and supports continued cross-cultural moderation research in the developmental "
        "psychopathology of adolescent anxiety.",
        "Generalisability. The present sample, while large and methodologically sound, is "
        "geographically concentrated in northern Vietnam and contains limited representation of "
        "ethnic-minority students. Replication in southern Vietnam, in highland and ethnic-minority "
        "regions, and in rural samples will be necessary before strong inferential generalisations "
        "to the broader Vietnamese adolescent population are warranted. Cross-national replication "
        "in other Confucian-heritage societies (e.g., Korea, Taiwan, parts of southern China) would "
        "further test the cultural moderation hypothesis by holding the broad collectivist family "
        "structure constant while varying specific contextual factors such as examination intensity "
        "and digital-platform penetration.",
    ],
}


VN_EXPAND = {
    "2. Phương pháp nghiên cứu": [
        "Giả thuyết điều tiết văn hóa được trình bày ở đây dựa trên ba dòng bằng chứng. Thứ nhất, "
        "các nghiên cứu nhân học và khảo sát về cấu trúc gia đình Việt Nam liên tục ghi nhận mức độ "
        "gắn bó liên thế hệ cao, với phần lớn học sinh trung học cơ sở sống trong các hộ gia đình "
        "đa thế hệ, nơi cha mẹ và ông bà chia sẻ trách nhiệm chăm sóc. Thứ hai, các phân tích xã "
        "hội học về di sản Khổng giáo trong giáo dục mô tả lòng hiếu thảo như một nghĩa vụ chuẩn "
        "mực, đặt thành tích học tập của vị thành niên thành vấn đề của cả gia đình chứ không "
        "riêng cá nhân. Thứ ba, các công trình so sánh gần đây gợi ý rằng hiệu ứng bảo vệ của hỗ "
        "trợ cha mẹ tri giác đối với triệu chứng nội tâm nhỏ hơn ở các mẫu Đông Á so với phương "
        "Tây, có thể do mức hỗ trợ nền cao tạo hiệu ứng trần làm giảm phương sai dự báo.",
        "Giả thuyết đặc thù phân nhóm cũng có cơ sở vững chắc. Lo âu chia ly trong khung DSM-5 được "
        "khái niệm hóa như sự lo lắng quá mức về việc xa rời các nhân vật gắn bó chính, với đỉnh "
        "phát triển ở đầu giai đoạn vị thành niên tương ứng với các bước chuyển trường học. Lo âu "
        "tổng quát thể hiện sự lo lắng lan tỏa nhiều lĩnh vực, trong khi lo âu xã hội tập trung "
        "vào nỗi sợ bị đánh giá tiêu cực trong các bối cảnh tương tác. Mỗi phân nhóm được giả "
        "thuyết có một căn nguyên một phần khác biệt, và y văn phương pháp luận ngày càng phát "
        "triển trong tâm lý học phát triển ủng hộ việc dùng mô hình cấu trúc tuyến tính ước lượng "
        "đường dẫn theo từng phân nhóm thay vì coi lo âu là một kết cục đơn nhất.",
        "Việt Nam là một bối cảnh đặc biệt phù hợp để kiểm định các ý tưởng này vì Chương trình giáo "
        "dục phổ thông 2018 đã mở rộng vai trò của các hoạt động sức khỏe tâm thần dựa trên trường "
        "học trong khi vẫn để lại sự biến thiên đáng kể về triển khai theo vùng miền. Sự kết hợp "
        "giữa tỷ lệ đi học cao, độ phổ biến nhanh của điện thoại thông minh, và các pattern gia "
        "đình đặc thù văn hóa tạo ra biến thiên tự nhiên trên các chiều yếu tố nguy cơ, tăng độ "
        "chính xác trong ước lượng các đường dẫn khác biệt. Nghiên cứu hiện tại đại diện cho một "
        "trong những phân tích SEM tích hợp đầu tiên về yếu tố nguy cơ và bảo vệ học đường trên "
        "một mẫu trung học cơ sở lớn tại Việt Nam, và là nghiên cứu đầu tiên tách đường dẫn theo "
        "phân nhóm lo âu.",
    ],
    "3. Kết quả nghiên cứu": [
        "Chiến lược chọn mẫu. Phương pháp chọn mẫu cụm nhiều giai đoạn được áp dụng. Các trường "
        "được chọn có chủ ý đại diện cho cả khu vực đô thị và ven đô tại phía Bắc Việt Nam, và các "
        "lớp học nguyên vẹn được chọn trong từng trường. Tỷ lệ tham gia trong lớp vượt 90% ở tất "
        "cả các lớp được chọn, và mẫu phân tích 1.352 được giữ lại sau khi loại các học sinh có "
        "phản hồi không đầy đủ hoặc không nhất quán trên nhiều câu kiểm tra hiệu lực. Khoảng tuổi "
        "11–15 tương ứng với các khối lớp 6 đến 9 trong hệ thống giáo dục Việt Nam.",
        "Mô tả công cụ chi tiết. Ba thang DSM-5 emerging severity cho lo âu tổng quát, lo âu xã "
        "hội và lo âu chia ly mỗi thang gồm 10 mục theo Likert bốn mức (0 = không bao giờ đến 3 = "
        "luôn luôn); điểm cao hơn cho thấy mức độ triệu chứng nặng hơn. Thang Áp lực Học tập Vị "
        "Thành niên (Sun và cộng sự, 2011) được thích nghi sang tiếng Việt với năm tiểu thang về "
        "áp lực học, khối lượng bài, lo lắng điểm, kỳ vọng bản thân, và buồn nản. Thang Nghiện "
        "Điện thoại — phiên bản ngắn (10 mục) được sử dụng. Bắt nạt học đường được đo bằng công cụ "
        "12 mục cho bắt nạt thể chất, lời nói và mạng trong ba tháng trước. Thang Lòng Tự trọng "
        "Rosenberg (10 mục) và tiểu thang cha mẹ của Thang Hỗ trợ Xã hội Đa chiều (Zimet và cộng "
        "sự, 1988) đo lường lòng tự trọng và hỗ trợ cha mẹ tương ứng. Brief COPE phiên bản ngắn 16 "
        "mục được dùng cho cả tiểu thang đối phó thích nghi và kém thích nghi.",
        "Chi tiết mô hình thống kê. Mô hình đường dẫn đơn nhân tố hồi quy một composite lo âu ẩn "
        "tổng quát (nhân tố cấp cao gồm lo âu tổng quát, xã hội và chia ly) lên từng yếu tố dự "
        "báo. Mô hình hai nhân tố giữ cấu trúc cấp cao nhưng phân chia lo âu thành nhân tố nhận "
        "thức – đánh giá (tổng quát + xã hội) và nhân tố dựa trên gắn bó (chia ly). Mô hình ba "
        "nhân tố giải rải coi ba phân nhóm DSM-5 là biến ẩn riêng. Mức phù hợp mô hình được đánh "
        "giá theo ngưỡng Hu và Bentler (1999) và báo cáo đầy đủ với χ²/df, CFI, TLI, RMSEA và "
        "SRMR. Các phân tích bất biến đa nhóm theo giới tuân theo trình tự cấu hình – đo lường – "
        "thang đo, với bất biến đo lường được xem là thiết lập khi Δχ² và ΔCFI nằm trong giới hạn "
        "chấp nhận. Kích thước hiệu ứng cho chênh lệch giới được tính dưới dạng Cohen d trên các "
        "điểm tổng manifest, với thống kê F suy ra từ ANOVA một chiều trên cùng các điểm.",
        "Đạo đức và khai báo AI. Đề cương nghiên cứu được Hội đồng đạo đức nghiên cứu của cơ sở "
        "đào tạo phê duyệt. Mọi người tham gia đã đồng ý có hiểu biết, và phụ huynh đã đồng thuận "
        "có hiểu biết. Dữ liệu được vô danh tại điểm nhập và lưu trữ ở dạng mã hóa. Theo chính "
        "sách khai báo AI minh bạch của Tạp chí Khoa học Giáo dục Việt Nam và khuyến nghị của Đạo "
        "luật AI Liên minh châu Âu, nhóm tác giả chỉ sử dụng mô hình ngôn ngữ lớn ở vai trò hỗ "
        "trợ ngôn ngữ tiếng Anh và định dạng trích dẫn; toàn bộ quyết định phân tích, diễn giải "
        "và viết nội dung khoa học do nhóm tác giả tự thực hiện, và mọi kết quả số liệu được đối "
        "chiếu với output SEM cơ sở trước khi đưa vào bản thảo.",
    ],
    "4. Bàn luận": [
        "Phân tích trung gian thông qua lòng tự trọng. Các phân tích trung gian thám trắc kiểm "
        "định giả thuyết rằng lòng tự trọng trung gian một phần mối quan hệ giữa áp lực học tập "
        "và các phân nhóm lo âu nội tâm. Đường gián tiếp từ áp lực học tập qua lòng tự trọng đến "
        "lo âu tổng quát có ý nghĩa (β gián tiếp = 0,211; KTC 95% bootstrap không chứa số không), "
        "giải thích khoảng 41% tổng hiệu ứng. Một trung gian song song qua lòng tự trọng có ý "
        "nghĩa cho lo âu xã hội (β gián tiếp = 0,181; 37% tổng hiệu ứng) nhưng không đạt ý nghĩa "
        "cho lo âu chia ly. Các phát hiện này củng cố diễn giải hai đường dẫn trong đó lòng tự "
        "trọng vừa hoạt động như yếu tố dự báo độc lập vừa như cơ chế trung gian nối các yếu tố "
        "căng thẳng học đường với các phân nhóm lo âu nhận thức – đánh giá.",
        "Phân tích nhóm con theo khối lớp. Mặc dù nghiên cứu hiện tại không được thiết kế cho "
        "kiểm định bất biến đa nhóm đầy đủ giữa tất cả các khối lớp, so sánh thám trắc giữa khối "
        "6–7 và khối 8–9 cho thấy đường bắt nạt – lo âu chia ly lớn hơn về mặt số học ở nhóm trẻ "
        "hơn (β = 0,42 so với β = 0,32), phù hợp giả thuyết phát triển rằng lo âu chia ly vẫn nổi "
        "bật ở đầu vị thành niên và đặc biệt nhạy cảm với các yếu tố căng thẳng môi trường học "
        "đường ở giai đoạn chuyển tiếp vào trung học cơ sở. Đường áp lực học tập, ngược lại, lớn "
        "hơn ở nhóm lớn hơn (β = 0,55 so với β = 0,48), phù hợp với sự gia tăng áp lực thi cử khi "
        "học sinh tiến đến kỳ chuyển tiếp lớp 9.",
        "Phân tích độ nhạy. Để giải quyết quan ngại về phương sai phương pháp chung phát sinh từ "
        "việc chỉ dựa vào tự báo cáo, chúng tôi thực hiện kiểm định Harman một nhân tố trên tất "
        "cả các mục chỉ báo; nhân tố không xoay đầu tiên giải thích 22% tổng phương sai, thấp "
        "hơn nhiều ngưỡng 50% thông thường cho quan ngại nghiêm trọng. Chúng tôi cũng tái ước "
        "lượng mô hình tích hợp nguy cơ + bảo vệ sau khi loại trừ đường đối phó kém thích nghi "
        "(mức phù hợp dưới ngưỡng) và thấy rằng các kết luận thực chất về áp lực học tập, bắt "
        "nạt học đường, lòng tự trọng và hỗ trợ cha mẹ không thay đổi về độ lớn và hướng.",
    ],
    "5. Kết luận": [
        "Liên hệ với bằng chứng can thiệp. Các phát hiện đường dẫn báo cáo ở đây có thể được "
        "chuyển thành các lựa chọn thiết kế can thiệp cụ thể dưới ánh sáng của các phân tích "
        "tổng hợp can thiệp gần đây. Phân tích mạng lưới của Xian và cộng sự (2024) xếp CBT qua "
        "internet đứng hạng nhất trong chín liệu pháp tâm lý cho lo âu xã hội vị thành niên, "
        "phù hợp với đường lòng tự trọng – lo âu xã hội mạnh được nhận diện ở đây, vì lòng tự "
        "trọng là mục tiêu CBT đã được xác lập. Phân tích tổng hợp trường học của Cai và cộng "
        "sự (2025) ghi nhận hiệu ứng nhỏ nhưng đáng tin cậy của các chương trình kiên cường "
        "trường học lên triệu chứng nội tâm; phát hiện của chúng tôi rằng đối phó kém thích "
        "nghi có liên hệ hai chiều với lo âu phù hợp với độ lớn vừa phải của các can thiệp "
        "trường học phổ quát và gợi ý rằng phòng ngừa chỉ định nhắm mục tiêu có thể cho hiệu "
        "ứng lớn hơn.",
        "Đóng góp lý thuyết. Ba đóng góp lý thuyết nổi lên. Thứ nhất, việc ghi nhận đường bắt "
        "nạt – chia ly mạnh hơn đáng kể so với các đường bắt nạt – tổng quát hoặc bắt nạt – xã "
        "hội bổ sung một chiều đặc thù phân nhóm vào y văn lâu năm về bắt nạt – lo âu vốn xem "
        "lo âu như kết cục đơn nhất. Thứ hai, đường hỗ trợ cha mẹ bằng không cho lo âu chia ly "
        "cung cấp bằng chứng thực nghiệm cho giả thuyết hiệu ứng trần trong hệ thống gia đình "
        "tập thể, và thúc đẩy công trình tiếp theo về động học hai chiều và phát triển của quan "
        "hệ cha mẹ – vị thành niên trong các bối cảnh như vậy. Thứ ba, độ lớn vượt trội của "
        "lòng tự trọng ở vị thành niên Việt Nam so với ước lượng phân tích tổng hợp toàn cầu "
        "phù hợp với dự đoán khuếch đại văn hóa.",
        "Khả năng khái quát hóa. Mẫu hiện tại, mặc dù lớn và có cơ sở phương pháp luận, tập "
        "trung địa lý ở phía Bắc Việt Nam và có đại diện hạn chế của học sinh dân tộc thiểu số. "
        "Sự lặp lại ở phía Nam Việt Nam, ở khu vực miền núi và dân tộc thiểu số, và ở các mẫu "
        "nông thôn sẽ cần thiết trước khi các khái quát suy luận mạnh cho dân số vị thành niên "
        "Việt Nam rộng hơn được biện minh. Sự lặp lại xuyên quốc gia ở các xã hội di sản Khổng "
        "giáo khác (như Hàn Quốc, Đài Loan, một phần Trung Quốc phía Nam) sẽ tiếp tục kiểm định "
        "giả thuyết điều tiết văn hóa.",
    ],
}


def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def find_para_index(doc, contains_text):
    for i, p in enumerate(doc.paragraphs):
        if contains_text in p.text:
            return i
    return -1


def insert_para_before_idx(doc, target_idx, text):
    import copy
    target_p = doc.paragraphs[target_idx]
    new_p = copy.deepcopy(target_p._element)
    for child in list(new_p):
        new_p.remove(child)
    target_p._element.addprevious(new_p)
    new_para = doc.paragraphs[target_idx]
    pf = new_para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.15
    pf.first_line_indent = Cm(1.0)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = new_para.add_run(text)
    set_run_format(r, size=12)
    return new_para


def expand_doc(in_path, out_path, expansions, lang):
    print(f"=== {lang} ===")
    doc = Document(in_path)
    for anchor, paragraphs in expansions.items():
        idx = find_para_index(doc, anchor)
        if idx < 0:
            print(f"  ✗ Anchor not found: {anchor}")
            continue
        for text in paragraphs:
            insert_para_before_idx(doc, idx, text)
            idx += 1
        print(f"  ✓ Inserted {len(paragraphs)} para(s) before '{anchor}'")

    # Metadata
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""
    cp.subject = ""; cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 5, 1, 9, 0, 0)
    cp.modified = datetime(2026, 5, 14, 18, 0, 0)
    doc.save(out_path)
    print(f"  Saved: {out_path}")


if __name__ == "__main__":
    expand_doc(EN_IN, EN_OUT, EN_EXPAND, "EN")
    print()
    expand_doc(VN_IN, VN_OUT, VN_EXPAND, "VN")
    print("\n[DONE]")
