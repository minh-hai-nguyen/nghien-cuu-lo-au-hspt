# -*- coding: utf-8 -*-
"""Batch phase 2: Update 134 file residual + safety filters.
- Skip bai-bao-khgdvn/* (Q2.5 cam)
- Skip *_BACKUP_*, *_BEFORE_*
- Skip _Archive/, _to_send_to_other_machine/
01/06/2026."""
import os, sys, io, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

REPL = [
    ('LAN TỎA', 'TỔNG QUÁT'),
    ('LAN TOẢ', 'TỔNG QUÁT'),
    ('Lan tỏa', 'Tổng quát'),
    ('Lan toả', 'Tổng quát'),
    ('lan tỏa', 'tổng quát'),
    ('lan toả', 'tổng quát'),
    ('RLLALT', 'RLLATQ'),
]


def process_file(full_path):
    try:
        d = Document(full_path)
    except Exception as e:
        return {'status': f'ERR: {e}', 'count': 0}

    total = 0

    def fix_para(p):
        nonlocal total
        # run-level first
        for r in p.runs:
            for old, new in REPL:
                if old in r.text:
                    r.text = r.text.replace(old, new)
                    total += 1
        # if text split across runs, rebuild
        if any(o in p.text for o, _ in REPL):
            if p.runs:
                full = p.text
                new_full = full
                for old, new in REPL:
                    new_full = new_full.replace(old, new)
                if new_full != full:
                    first = p.runs[0]
                    fn, fs, fb, fi = first.font.name, first.font.size, first.bold, first.italic
                    for r in list(p.runs):
                        r._element.getparent().remove(r._element)
                    nr = p.add_run(new_full)
                    nr.font.name = fn
                    if fs: nr.font.size = fs
                    nr.bold = fb; nr.italic = fi
                    total += 1

    for p in d.paragraphs:
        fix_para(p)
    for tb in d.tables:
        for row in tb.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    fix_para(p)

    if total > 0:
        # Clean metadata
        cp = d.core_properties
        cp.author = ''; cp.title = ''; cp.subject = ''
        cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
        d.save(full_path)

    return {'status': 'OK', 'count': total}


# Find all docx with residual
SKIP_DIRS = ['bai-bao-khgdvn', '_Archive', '_to_send_to_other_machine',
             '.git', 'node_modules', '_workspace_inventory']
SKIP_PATTERNS = ['_BACKUP_', '_BEFORE_', 'BACKUP.docx', '_BACKUP.docx']

processed = 0
skipped = 0
total_replacements = 0
errors = []
files_updated = []

for root, dirs, files in os.walk(ROOT):
    # Skip dirs
    if any(s in root for s in SKIP_DIRS):
        dirs.clear()
        continue
    for f in files:
        if not f.lower().endswith('.docx'):
            continue
        # Skip backup patterns
        if any(p in f for p in SKIP_PATTERNS):
            skipped += 1
            continue
        full = os.path.join(root, f)
        # Quick check if file has residual
        try:
            d = Document(full)
            text = '\n'.join(p.text for p in d.paragraphs)[:50000]  # first 50K chars
            for tb in d.tables[:5]:
                for row in tb.rows[:20]:
                    for c in row.cells:
                        text += '\n' + c.text
            if not (any(o in text for o, _ in REPL)):
                continue  # no residual, skip
        except:
            continue

        # Process
        r = process_file(full)
        if r['count'] > 0:
            files_updated.append((full.replace(ROOT, '.'), r['count']))
            total_replacements += r['count']
            processed += 1
            if processed % 10 == 0:
                print(f'  Progress: {processed} files processed, '
                      f'{total_replacements} replacements...')
        if 'ERR' in r['status']:
            errors.append((full, r['status']))


print()
print('='*70)
print(f'PHASE 2 BATCH UPDATE COMPLETE')
print('='*70)
print(f'Files updated: {processed}')
print(f'Files skipped (backup pattern): {skipped}')
print(f'Total replacements: {total_replacements}')
print(f'Errors: {len(errors)}')

if files_updated:
    print()
    print('=== TOP UPDATED FILES ===')
    for f, c in sorted(files_updated, key=lambda x: -x[1])[:20]:
        print(f'  {c:>4} | {f}')

if errors:
    print()
    print('=== ERRORS ===')
    for f, e in errors[:10]:
        print(f'  {f} | {e}')
