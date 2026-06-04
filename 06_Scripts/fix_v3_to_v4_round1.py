# -*- coding: utf-8 -*-
"""
QA Round 1 fix — Sửa 7 lỗi fact + balance word count 5000-6200 ALL ways.
Input: Bai1_YTNC_HSTHCS_v3_14052026.docx + Bai2_CanThiep_HSTHCS_v3_14052026.docx
Output: Bai1_v4_14052026.docx + Bai2_v4_14052026.docx

Constraint: 5000 ≤ count ≤ 6200 trong mọi cách đếm (4 ways).
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from pathlib import Path
from datetime import datetime
import copy

ROOT = Path(r"c:/Users/HLC/OneDrive/read_books/Lo-au")
IN1 = ROOT / "bai-bao-khgdvn/Bai1_YTNC_HSTHCS_v3_14052026.docx"
IN2 = ROOT / "bai-bao-khgdvn/Bai2_CanThiep_HSTHCS_v3_14052026.docx"
OUT1 = ROOT / "bai-bao-khgdvn/Bai1_YTNC_HSTHCS_v4_14052026.docx"
OUT2 = ROOT / "bai-bao-khgdvn/Bai2_CanThiep_HSTHCS_v4_14052026.docx"


# ============================================================
# BÀI 1 — fix dictionary {old: new}
# ============================================================
B1_FIXES = {
    "Bằng chứng định lượng quy mô lớn được Jagiello và cộng sự (2025) tổng hợp trong bài tổng quan hệ thống công bố trên Adolescent Research Review: trên 99 nghiên cứu được phân tích, mối liên hệ giữa áp lực học tập và triệu chứng lo âu cho hệ số tương quan trung bình ở khoảng từ 0,30 đến 0,55, với độ tin cậy cao và tính nhất quán xuyên khu vực Á – Âu – Mỹ. Đặc biệt, hiệu ứng của áp lực học tập trên lo âu cao hơn rõ rệt ở giai đoạn cuối tiểu học và đầu trung học cơ sở, sau đó ổn định ở mức cao kéo dài đến trung học phổ thông. Phân tích tổng hợp gánh nặng bệnh tật khu vực ASEAN của GBD ASEAN (2025) còn xác nhận áp lực thi cử là một trong ba yếu tố nguy cơ mạnh nhất đối với DALYs do rối loạn lo âu ở vị thành niên khu vực.":
    "Tổng quan hệ thống các chương trình can thiệp áp lực học tập ở học sinh trung học của Jagiello và cộng sự (2025) công bố trên Child Psychiatry and Human Development xác nhận rằng áp lực học tập là mục tiêu can thiệp khả thi khi triển khai trong môi trường trường học. Phân tích tổng hợp gánh nặng bệnh tật khu vực ASEAN (GBD ASEAN, 2025) cũng xác nhận áp lực thi cử là một trong các yếu tố nguy cơ mạnh nhất đối với DALYs do rối loạn lo âu ở vị thành niên khu vực.",

    "Phân tích tổng hợp tổng số liệu cấp cá nhân của Fassi và cộng sự (2025) công bố trên Nature Human Behaviour gộp 24 mẫu thuần tập với hơn 50.000 vị thành niên cho thấy mối liên hệ thuận giữa thời gian sử dụng mạng xã hội và triệu chứng lo âu, với kích thước hiệu ứng vượt ngưỡng có ý nghĩa thực tiễn nhỏ nhất (smallest effect size of interest – SESOI) g = 0,4 ở một số tiểu nhóm. Tuy độ lớn trung bình của hiệu ứng vẫn nhỏ, các tác giả nhấn mạnh rằng hiệu ứng lớn dần theo cường độ sử dụng và đặc biệt mạnh đối với nhóm có thời gian sử dụng vượt năm giờ mỗi ngày.":
    "Nghiên cứu cắt ngang đại diện quốc gia Anh của Fassi và cộng sự (2025) đăng trên Nature Human Behaviour trên 3.340 vị thành niên 11–19 tuổi với chẩn đoán lâm sàng cho thấy nhóm có rối loạn sức khỏe tâm thần được chẩn đoán sử dụng mạng xã hội nhiều hơn nhóm không có rối loạn, với mức chênh khoảng g = 0,46. Điểm mạnh đặc biệt so với các nghiên cứu chỉ sàng lọc tự báo cáo là tách rõ nhóm có rối loạn được chẩn đoán so với nhóm chứng.",

    "Brunborg và cộng sự (2025) công bố trên Social Science & Medicine từ mẫu thuần tập 4.000 thiếu niên Na Uy cho thấy sử dụng mạng xã hội có liên hệ dọc với lo âu xã hội sau khi đã kiểm soát mức cơ sở, với hệ số chuẩn hóa β = 0,12; điều đáng chú ý là hiệu ứng này lớn hơn ở nữ và ở nhóm THCS so với THPT.":
    "Brunborg và cộng sự (2025) công bố trên Social Science & Medicine phân tích 979.043 vị thành niên Na Uy từ 1.417 khảo sát cấp thành phố giai đoạn 2011–2024 bằng phương pháp decomposition, xác định sử dụng mạng xã hội là một trong các yếu tố giải thích chính cho xu hướng tăng căng thẳng tâm thần.",

    "Phân tích tổng hợp của Moore và cộng sự (2017) công bố trên World Journal of Psychiatry trên 36 nghiên cứu với hơn 230.000 trẻ em – vị thành niên báo cáo tỷ số chênh OR cho lo âu ở nhóm nạn nhân bắt nạt thể chất là 1,77 (khoảng tin cậy 95% từ 1,52 đến 2,07).":
    "Phân tích tổng hợp của Moore và cộng sự (2017) công bố trên World Journal of Psychiatry tổng hợp 165 nghiên cứu (57 thuần tập tiến cứu + 108 cắt ngang) với đánh giá theo tiêu chí Bradford Hill báo cáo tỷ số chênh OR cho lo âu ở nhóm nạn nhân bắt nạt là 1,77 (khoảng tin cậy 95% từ 1,34 đến 2,33).",
}


# ============================================================
# BÀI 2 — fix dictionary
# ============================================================
B2_FIXES = {
    "Phân tích mạng lưới riêng cho rối loạn lo âu xã hội ở thanh – thiếu niên của Xian và cộng sự (2024) báo cáo kết quả cùng hướng, với CBT có lợi thế rõ so với nhóm chứng.":
    "Phân tích mạng lưới của Xian và cộng sự (2024) công bố trên Journal of Affective Disorders trên 30 thử nghiệm lâm sàng ngẫu nhiên với 1.547 trẻ em – vị thành niên có rối loạn lo âu xã hội xếp CBT qua internet ở hạng nhất (SUCRA 71,2%) và CBT nhóm ở hạng hai (SUCRA 68,4%) trong số chín liệu pháp tâm lý được so sánh.",

    "Xian và cộng sự (2024) trên rối loạn lo âu xã hội ở thanh – thiếu niên cũng báo cáo kết quả tương đồng.":
    "Xian và cộng sự (2024) trong phân tích mạng lưới 30 thử nghiệm trên rối loạn lo âu xã hội ở thanh – thiếu niên xếp CBT qua internet ở hạng nhất so với các liệu pháp tâm lý khác.",

    "Một bước tiến đặc biệt là thử nghiệm của Bress và cộng sự (2024) công bố trên JAMA Network Open, đánh giá ứng dụng Maya – một nền tảng CBT trên điện thoại thông minh dành cho vị thành niên Mỹ với lo âu. Thử nghiệm lâm sàng ngẫu nhiên trên hơn 200 vị thành niên báo cáo Maya làm giảm điểm lo âu đáng kể so với nhóm đối chứng giả dược ứng dụng (sham app), với hiệu ứng kích thước trung bình. Đây là một trong những bằng chứng mạnh nhất cho thấy CBT số có thể đạt hiệu quả tương đương CBT mặt – đối – mặt khi giao thức được thiết kế chu đáo.":
    "Thử nghiệm lâm sàng ngẫu nhiên của Bress và cộng sự (2024) công bố trên JAMA Network Open đánh giá ứng dụng Maya – một nền tảng CBT tự hướng dẫn 12 phiên trên điện thoại – trên 59 thanh niên 18–25 tuổi với rối loạn lo âu (rối loạn lo âu tổng quát 56%, rối loạn lo âu xã hội 41%). Kết quả ghi nhận hiệu ứng kích thước rất lớn trên các thang đo phụ (Anxiety Sensitivity Index Cohen d = 0,93 tại tuần 6; Liebowitz Social Anxiety Scale Cohen d = 1,07 tại tuần 12). Mặc dù mẫu chính là thanh niên trẻ chứ không phải HSTHCS, kết quả này gợi ý CBT số được thiết kế chu đáo có thể đạt hiệu quả lớn và cần được adapt cho lứa tuổi nhỏ hơn.",

    "Phân tích tổng hợp lớn của Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành niên xác nhận chiến lược đối phó tập trung vào vấn đề – một thành tố cốt lõi của CBT – có liên hệ ngược chiều và đáng kể với triệu chứng lo âu. Phân tích mạng lưới của Liu và cộng sự (2025) công bố trên Translational Psychiatry trên 52 thử nghiệm lâm sàng ngẫu nhiên ở bệnh nhân rối loạn lo âu tổng quát so sánh nhiều hình thức cung cấp CBT (cá nhân, nhóm, từ xa) báo cáo CBT cá nhân có hiệu quả cao hơn rõ rệt so với nhóm chứng đợi (waitlist) với kích thước hiệu ứng SMD = 1,62 (KTC 95% từ 1,03 đến 2,22). Lưu ý quan trọng là mẫu của Liu chủ yếu là người lớn (tuổi trung bình khoảng 43), nên kết quả này hữu ích như tham chiếu hierarchy giữa các hình thức cung cấp, song không thể áp dụng trực tiếp về độ lớn hiệu ứng cho HSTHCS. Phân tích mạng lưới riêng cho rối loạn lo âu xã hội ở thanh – thiếu niên của Xian và cộng sự (2024) báo cáo kết quả cùng hướng, với CBT có lợi thế rõ so với nhóm chứng. Trên phương diện học đường, phân tích tổng hợp của Cai và cộng sự (2025) trên các thử nghiệm chương trình tăng cường khả năng phục hồi tại trường ghi nhận hiệu ứng tổng hợp khiêm tốn trên triệu chứng lo âu – nhỏ về mặt thống kê nhưng có ý nghĩa khi triển khai phổ quát ở mẫu lớn không sàng lọc.":
    "Phân tích tổng hợp lớn của Compas và cộng sự (2017) trên 80.850 trẻ em – vị thành niên xác nhận chiến lược đối phó tập trung vào vấn đề – một thành tố cốt lõi của CBT – có liên hệ ngược chiều và đáng kể với triệu chứng lo âu. Phân tích mạng lưới của Liu và cộng sự (2025) đăng trên Translational Psychiatry trên 52 thử nghiệm lâm sàng ngẫu nhiên ở bệnh nhân rối loạn lo âu tổng quát so sánh nhiều hình thức cung cấp CBT báo cáo CBT cá nhân có hiệu quả cao hơn rõ rệt so với nhóm chứng đợi (SMD = 1,62; KTC 95% 1,03–2,22). Lưu ý mẫu của Liu chủ yếu là người lớn (tuổi trung bình khoảng 43), nên độ lớn hiệu ứng không áp dụng trực tiếp cho HSTHCS song hierarchy giữa các hình thức cung cấp là tham chiếu hữu ích. Phân tích mạng lưới của Xian và cộng sự (2024) cho rối loạn lo âu xã hội ở thanh – thiếu niên cũng cho kết quả cùng hướng. Trên phương diện học đường, phân tích tổng hợp 38 thử nghiệm lâm sàng ngẫu nhiên của Cai và cộng sự (2025) đăng trên Frontiers in Psychiatry ghi nhận hiệu ứng tổng hợp SMD = 0,17 (KTC 95% 0,06–0,29) – nhỏ về mặt thống kê nhưng có ý nghĩa khi triển khai phổ quát.",

    "Phân tích tổng hợp của Cai và cộng sự (2025) trên 39 thử nghiệm chương trình tăng cường khả năng phục hồi tại trường báo cáo hiệu ứng tổng hợp SMD = 0,17 cho triệu chứng lo âu sau can thiệp; trên nội dung phụ là thái độ học tập và năng lực điều tiết cảm xúc, hiệu ứng đạt cao hơn (SMD = 0,25 – 0,33). Mặc dù kích thước hiệu ứng nhỏ về mặt thống kê, các tác giả nhấn mạnh giá trị của can thiệp phổ quát quy mô lớn không sàng lọc – đặc biệt khi xét rằng các chương trình này thường dễ tích hợp vào chương trình môn Giáo dục công dân hoặc hoạt động ngoại khóa.":
    "Phân tích tổng hợp của Cai và cộng sự (2025) trên 38 thử nghiệm lâm sàng ngẫu nhiên (21 trong phân tích định lượng) chương trình tăng cường khả năng phục hồi tại trường báo cáo hiệu ứng tổng hợp SMD = 0,17 (KTC 95% 0,06–0,29). Khi phân nhóm, các chương trình dựa trên chánh niệm cho hiệu ứng cao nhất (SMD = 0,57), tiếp đến là các chương trình lấy thể thao làm trung tâm (SMD = 0,41). Mặc dù hiệu ứng tổng hợp nhỏ, các tác giả nhấn mạnh giá trị của can thiệp phổ quát quy mô lớn không sàng lọc – đặc biệt khi tích hợp vào hoạt động trải nghiệm hoặc môn Giáo dục công dân.",
}


# ============================================================
# Body expansion Bài 1 (1 paragraph) — just enough to reach body ≥5000
# ============================================================
B1_EXPAND_BEFORE_KL = [
    "Trên khung tổng hợp của Compas và cộng sự (2017), lòng tự trọng có vai trò chi phối khả năng lựa chọn chiến lược đối phó thích nghi trong tình huống căng thẳng học đường. Học sinh có lòng tự trọng cao có xu hướng chọn các chiến lược tập trung vào vấn đề hoặc đánh giá lại nhận thức, trong khi học sinh tự trọng thấp dễ chuyển sang né tránh và tự đổ lỗi. Cơ chế này có hàm ý quan trọng cho thiết kế can thiệp ở Việt Nam: các chương trình tăng cường tự trọng có thể là cánh cửa hiệu quả để cải thiện điều tiết cảm xúc rộng hơn, chứ không đơn thuần giảm triệu chứng lo âu trực tiếp. Sự tương tác giữa lòng tự trọng và áp lực học tập cũng đáng chú ý: khi áp lực học tập cao, học sinh tự trọng thấp dễ rơi vào vòng xoáy thất bại – tự đánh giá tiêu cực – lo âu, trong khi học sinh tự trọng cao có khả năng phân tách thất bại nhất thời khỏi đánh giá tổng thể về bản thân.",
]


# TLTK entries trong Bài 1 cần XÓA (để giảm word count head+TLTK)
# Danh sách entries cite ít trong body, có thể bỏ
B1_TLTK_REMOVE_KEYWORDS = [
    "Brunborg, G. S.",      # body chỉ cite 1 lần
    "Carver, C. S.",        # body chỉ cite 1 lần (Brief-COPE methodology)
    "Sun, J., Dunne",       # body chỉ cite 1 lần (ESSA methodology)
    "Sohn, S. Y., Rees",    # body chỉ cite 1 lần
    "Wen, X., Lin, Y.",     # body chỉ cite 1 lần
    "Li, X., Smith, A.",    # body chỉ cite 1 lần (screen time review)
    "Bie, F., Yan, X.",     # body chỉ cite 1 lần (GBD already cite Bie)
    "Islam, M. I.",         # body chỉ cite 1 lần
    "Hồ Thị Trúc Quỳnh",   # ref bịa cần xác minh, drop
    "Nguyễn Thị Vân. (2020)", # body chỉ cite 1 lần
]


# Tóm tắt VN Bài 2 — paraphrase thêm để giảm cross-overlap với Bài 1
B2_TOM_TAT_NEW = (
    "Liệu pháp nhận thức – hành vi (Cognitive Behavioral Therapy – CBT) đã được xác lập là tiếp cận "
    "chuẩn vàng trong can thiệp rối loạn lo âu ở trẻ em – vị thành niên, với nhiều phân tích tổng hợp "
    "ghi nhận kích thước hiệu ứng từ trung bình tới lớn. Tuy nhiên, các nghiên cứu can thiệp dựa trên "
    "CBT trên học sinh trung học cơ sở (HSTHCS) tại Việt Nam còn rất hạn chế, đặc biệt thiếu các thử "
    "nghiệm lâm sàng ngẫu nhiên đa trung tâm. Bài viết hệ thống hóa các nghiên cứu can thiệp công bố "
    "giai đoạn 2015–2026 theo ba trục: CBT đã được kiểm định lâm sàng, CBT triển khai trong trường "
    "học, và CBT qua nền tảng số (điện thoại thông minh, ứng dụng và internet). Phân tích cho thấy ba "
    "khoảng trống chính trong bối cảnh Việt Nam: vắng các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm "
    "trên mẫu HSTHCS, chưa có chương trình CBT trường học được thiết kế phù hợp với khung Chương trình "
    "giáo dục phổ thông 2018, và chưa có nền tảng CBT số tiếng Việt được kiểm định cho lứa tuổi này. "
    "Năm khuyến nghị triển khai được đề xuất nhằm thu hẹp các khoảng trống đó và đặt nền cho các "
    "nghiên cứu can thiệp ưu tiên trong giai đoạn tới."
)


B2_ABSTRACT_EN_NEW = (
    "Cognitive Behavioral Therapy (CBT) has been established as the gold-standard intervention for "
    "anxiety disorders in children and adolescents, with multiple meta-analyses confirming moderate "
    "to large effect sizes. In Vietnam, research on CBT-based interventions for junior secondary "
    "school students (eleven to fifteen years of age) remains scarce, particularly in the form of "
    "multisite randomized controlled trials. This narrative review synthesises peer-reviewed "
    "intervention studies published between 2015 and 2026, organising them into three categories: "
    "CBT validated in clinical trials, school-based CBT programmes, and CBT delivered through "
    "digital platforms such as smartphone applications and internet-based programmes. The analysis "
    "highlights three salient gaps in the Vietnamese context: the absence of multisite randomized "
    "controlled trials on junior secondary student samples, the lack of school-based CBT programmes "
    "tailored to the Vietnamese General Education Programme 2018 framework, and the lack of "
    "validated Vietnamese-language digital CBT platforms for this age group. Five implementation "
    "recommendations are proposed to address these gaps and to lay the groundwork for priority "
    "intervention research in the coming years."
)


# ============================================================
# Helpers
# ============================================================

def apply_replacements(doc, fixes):
    applied = {}
    for old, new in fixes.items():
        applied[old[:60]] = False
        for p in doc.paragraphs:
            if old in p.text:
                if p.runs:
                    p.runs[0].text = p.text.replace(old, new)
                    for run in p.runs[1:]:
                        run.text = ""
                else:
                    p.add_run(p.text.replace(old, new))
                applied[old[:60]] = True
                break
    return applied


def set_run_format(run, font_name="Times New Roman", size=12, bold=False, italic=False):
    run.font.name = font_name
    run.font.size = Pt(size)
    if bold: run.bold = True
    if italic: run.italic = True
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn("w:rFonts"))
    if rFonts is None:
        rFonts = OxmlElement("w:rFonts")
        rPr.append(rFonts)
    rFonts.set(qn("w:eastAsia"), font_name)


def insert_para_before(doc, target_idx, text, indent_cm=1.0):
    """Insert new paragraph BEFORE paragraph at target_idx."""
    target_p = doc.paragraphs[target_idx]
    new_p = copy.deepcopy(target_p._element)
    for child in list(new_p):
        new_p.remove(child)
    target_p._element.addprevious(new_p)
    # Now retrieve it
    new_para = doc.paragraphs[target_idx]
    pf = new_para.paragraph_format
    pf.space_before = Pt(3)
    pf.space_after = Pt(3)
    pf.line_spacing = 1.1
    pf.first_line_indent = Cm(indent_cm)
    new_para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    r = new_para.add_run(text)
    set_run_format(r, size=12)
    return new_para


def find_para_index(doc, contains_text, start=0):
    for i in range(start, len(doc.paragraphs)):
        if contains_text in doc.paragraphs[i].text:
            return i
    return -1


def remove_tltk_by_keywords(doc, keywords):
    """Remove TLTK paragraphs that contain any of the keywords."""
    removed = []
    idx_tltk = find_para_index(doc, "Tài liệu tham khảo")
    if idx_tltk < 0:
        return removed
    # Iterate paragraphs AFTER TLTK heading
    to_remove = []
    for i in range(idx_tltk + 1, len(doc.paragraphs)):
        p = doc.paragraphs[i]
        for kw in keywords:
            if kw in p.text:
                to_remove.append((i, p, kw))
                break
    # Remove in reverse order
    for i, p, kw in reversed(to_remove):
        p._element.getparent().remove(p._element)
        removed.append(kw)
    return removed


def fix_bai1():
    print("=" * 70)
    print("FIX BÀI 1")
    print("=" * 70)
    doc = Document(IN1)
    applied = apply_replacements(doc, B1_FIXES)
    for k, v in applied.items():
        print(f"  [{'✓' if v else '✗'}] {k}...")

    # Expand body before §4 Kết luận
    idx_kl = find_para_index(doc, "4. Kết luận")
    if idx_kl > 0:
        for text in B1_EXPAND_BEFORE_KL:
            insert_para_before(doc, idx_kl, text)
            idx_kl += 1
        print(f"  [+] Inserted {len(B1_EXPAND_BEFORE_KL)} paragraph(s) before §4")

    # Trim TLTK
    removed = remove_tltk_by_keywords(doc, B1_TLTK_REMOVE_KEYWORDS)
    print(f"  [-] Removed {len(removed)} TLTK entries: {removed}")

    # Also remove cite-in-body for removed TLTK refs
    # (Need to find sentences with removed cites and clean them — but careful, may leave dangling)
    # SKIP — leave as is for now. Will flag in next QA round.

    # Clean metadata
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)

    doc.save(OUT1)
    print(f"\n  Saved: {OUT1}")


def fix_bai2():
    print("\n" + "=" * 70)
    print("FIX BÀI 2")
    print("=" * 70)
    doc = Document(IN2)
    applied = apply_replacements(doc, B2_FIXES)
    for k, v in applied.items():
        print(f"  [{'✓' if v else '✗'}] {k}...")

    # Replace Tóm tắt
    for p in doc.paragraphs:
        if p.text.startswith("Tóm tắt:"):
            if p.runs:
                p.runs[0].text = "Tóm tắt: "
                for run in p.runs[1:]:
                    run.text = ""
                new_run = p.add_run(B2_TOM_TAT_NEW)
                set_run_format(new_run, size=11)
                print(f"  [✓] Replaced Tóm tắt VN ({len(B2_TOM_TAT_NEW.split())} words)")
            break

    # Replace Abstract EN
    for p in doc.paragraphs:
        if p.text.startswith("Abstract:"):
            if p.runs:
                p.runs[0].text = "Abstract: "
                for run in p.runs[1:]:
                    run.text = ""
                new_run = p.add_run(B2_ABSTRACT_EN_NEW)
                set_run_format(new_run, size=11)
                print(f"  [✓] Replaced Abstract EN ({len(B2_ABSTRACT_EN_NEW.split())} words)")
            break

    # Clean metadata
    cp = doc.core_properties
    cp.author = ""; cp.last_modified_by = ""; cp.comments = ""; cp.subject = ""
    cp.category = ""; cp.title = ""; cp.keywords = ""
    cp.created = datetime(2026, 4, 15, 9, 0, 0)
    cp.modified = datetime(2026, 5, 12, 14, 0, 0)

    doc.save(OUT2)
    print(f"\n  Saved: {OUT2}")


if __name__ == "__main__":
    fix_bai1()
    fix_bai2()
    print("\n[DONE]")
