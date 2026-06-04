# -*- coding: utf-8 -*-
"""Sinh 2 file ho tro hoan thien luan an:
  File 05: Ly do chon de tai (~500 tu)
  File 06: Kien nghi + Huong nghien cuu (~800 tu)
Material: trich tu Tong quan v2 30-trang + Bai 1 + Bai 2.
Ngay 26/05/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUTDIR = os.path.join(ROOT, 'Luận án TS')

# ============================================================
# HELPERS
# ============================================================
def shade(cell, color):
    sh = OxmlElement('w:shd'); sh.set(qn('w:fill'), color); sh.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def colw(cell, cm):
    tcPr = cell._tc.get_or_add_tcPr()
    w = OxmlElement('w:tcW'); w.set(qn('w:w'), str(int(cm*567))); w.set(qn('w:type'), 'dxa')
    for old in tcPr.findall(qn('w:tcW')): tcPr.remove(old)
    tcPr.append(w)
def set_grid(t, widths):
    tblPr = t._tbl.find(qn('w:tblPr'))
    layout = OxmlElement('w:tblLayout'); layout.set(qn('w:type'), 'fixed')
    if tblPr is not None:
        for old in tblPr.findall(qn('w:tblLayout')): tblPr.remove(old)
        tblPr.append(layout)
    tg = t._tbl.find(qn('w:tblGrid'))
    if tg is not None: t._tbl.remove(tg)
    tg = OxmlElement('w:tblGrid')
    for w in widths:
        gc = OxmlElement('w:gridCol'); gc.set(qn('w:w'), str(int(w*567))); tg.append(gc)
    first_tr = t._tbl.findall('.//' + qn('w:tr'))[0]
    t._tbl.insert(list(t._tbl).index(first_tr), tg)

def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.5)
    return doc

def H(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0, 0, 0)

def P(doc, text, indent=True, italic=False, bold=False):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if indent:
        p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    r.italic = italic; r.bold = bold

# ============================================================
# FILE 05: LY DO CHON DE TAI (~500 tu, +/- 30%)
# ============================================================
doc5 = doc_init()
H(doc5, 'LÝ DO CHỌN ĐỀ TÀI', 1)
P(doc5, '(Bản đề xuất viết lại — trích tích hợp các khoảng trống nghiên cứu đã trình bày trong Tổng quan và hai bài đính kèm)', indent=False, italic=True)
doc5.add_paragraph()

P(doc5, 'Rối loạn lo âu là một trong những vấn đề sức khỏe tâm thần phổ biến nhất ở vị thành niên trên toàn cầu. Tại Việt Nam, khảo sát đại diện quốc gia V-NAMHS (Viện Xã hội học và cộng sự, 2022) trên 5.996 cặp cha mẹ — vị thành niên cho thấy rối loạn lo âu là rối loạn tâm thần phổ biến nhất ở lứa tuổi 10-17, với tỷ lệ đạt ngưỡng chẩn đoán DSM-5 là 2,3% và tỷ lệ có triệu chứng cận lâm sàng kèm suy giảm chức năng lên tới 21,7%. Các khảo sát học đường gần đây — Phạm Thị Thu Hoa và cộng sự (2024) trên 3.910 học sinh trung học phổ thông Hà Nội, Hoàng Trung Học và Nguyễn Thùy Dung (2025) trên 8.389 thanh thiếu niên ở 6 tỉnh, Trúc Thanh Thái và cộng sự (2026) trên 2.631 học sinh tại Thành phố Hồ Chí Minh — đều ghi nhận tỷ lệ triệu chứng lo âu cao, trong nhiều trường hợp vượt 40%. Đại dịch COVID-19 đã làm gia tăng đáng kể gánh nặng này, đồng thời để lại những thay đổi dài hạn về thói quen sử dụng thiết bị số và mạng xã hội ở lứa tuổi vị thành niên.')

P(doc5, 'Học sinh trung học cơ sở (11-15 tuổi) là lứa tuổi có ý nghĩa đặc biệt trong việc nghiên cứu rối loạn lo âu vì hai lý do. Thứ nhất, đây là giai đoạn khởi phát đặc trưng của nhiều dạng rối loạn lo âu — đặc biệt là rối loạn lo âu xã hội, vốn có tuổi khởi phát trung bình 11-13. Thứ hai, lứa tuổi này đứng giữa chuyển tiếp từ tiểu học lên trung học cơ sở và chuẩn bị thi vào lớp 10 — hai mốc gây áp lực học tập lớn trong hệ thống giáo dục Việt Nam. Nghiên cứu sớm trên lứa tuổi này không chỉ giúp can thiệp trước khi rối loạn tiến triển thành rối loạn lâm sàng kéo dài, mà còn tận dụng được hệ thống trường học làm kênh tiếp cận hiệu quả nhất về chi phí và quy mô.')

P(doc5, 'Mặc dù sự gia tăng ghi nhận về tỷ lệ rối loạn lo âu, nghiên cứu trong nước về chủ đề này vẫn còn sáu khoảng trống quan trọng. Thứ nhất, chưa có bộ thang đo lo âu cho trẻ em — vị thành niên được thẩm định tâm trắc toàn diện cho dân số Việt Nam, dẫn đến biên độ ước tính tỷ lệ rất rộng (từ 2,3% theo chẩn đoán đầy đủ đến trên 50% theo thang sàng lọc). Thứ hai, phần lớn nghiên cứu tập trung vào lo âu tổng quát qua thang GAD-7 hoặc DASS-21, trong khi lo âu xã hội và lo âu chia ly hầu như chưa được khảo sát chuyên biệt. Thứ ba, yếu tố bảo vệ — bao gồm lòng tự trọng, gắn bó trường học, hỗ trợ của cha mẹ và bạn bè, biện pháp ứng phó — chưa được nghiên cứu đồng thời và tích hợp với yếu tố nguy cơ trong cùng thiết kế. Thứ tư, thiết kế dọc và thử nghiệm lâm sàng ngẫu nhiên đa trung tâm còn rất ít. Thứ năm, mẫu nghiên cứu tập trung ở Hà Nội và Thành phố Hồ Chí Minh; khu vực nông thôn, miền núi và dân tộc thiểu số còn ít được khảo sát chuyên biệt. Thứ sáu, các chương trình can thiệp nhận thức — hành vi bản địa được kiểm định trên học sinh trung học cơ sở Việt Nam còn vắng bóng. Phân tích chi tiết các khoảng trống này được trình bày trong Tổng quan của luận án cùng hai bài tổng hợp chuyên sâu của tác giả đính kèm — Bài 1 về Yếu tố nguy cơ và Bài 2 về Khoảng trống can thiệp.')

P(doc5, 'Trên cơ sở các khoảng trống trên, luận án đặt mục tiêu nghiên cứu thực trạng rối loạn lo âu ở học sinh trung học cơ sở Hà Nội trên mẫu 1.352 học sinh tại hai trường có đặc thù xã hội — kinh tế khác nhau (Nhật Tân và Tây Mỗ), với thiết kế đo lường đồng thời các yếu tố nguy cơ và yếu tố bảo vệ chính, kết hợp khảo sát định lượng và phỏng vấn định tính. Trên cơ sở dữ liệu thực trạng, luận án đề xuất khung chương trình tập huấn phòng ngừa rối loạn lo âu cho học sinh trung học cơ sở, đặt nền cho các nghiên cứu can thiệp tiếp theo. Luận án vì vậy đóng góp một bước cụ thể vào việc thu hẹp ba trong sáu khoảng trống nêu trên — đo lường tích hợp yếu tố nguy cơ và bảo vệ, chú trọng lứa tuổi trung học cơ sở, và đặt cơ sở cho can thiệp bản địa hoá.')

out5 = os.path.join(OUTDIR, '05_LyDoChonDeTai_v1_26052026.docx')
doc5.save(out5)

# ============================================================
# FILE 06: KIEN NGHI + HUONG NGHIEN CUU (~800 tu)
# ============================================================
doc6 = doc_init()
H(doc6, 'KIẾN NGHỊ VÀ HƯỚNG NGHIÊN CỨU TIẾP THEO', 1)
P(doc6, '(Bản đề xuất tích hợp từ Tổng quan luận án, Bài 1 — Yếu tố nguy cơ và Bài 2 — Khoảng trống can thiệp)', indent=False, italic=True)
doc6.add_paragraph()

P(doc6, 'Trên cơ sở các phát hiện thực trạng và phân tích các khoảng trống nghiên cứu được trình bày xuyên suốt luận án, tác giả xin đề xuất sáu nhóm kiến nghị nhằm thu hẹp các khoảng trống và định hướng cho các nghiên cứu tiếp theo trong giai đoạn 2026-2030.')

H(doc6, '1. Về công cụ đo lường rối loạn lo âu cho học sinh trung học cơ sở Việt Nam', 2)
P(doc6, 'Cần xây dựng và thẩm định tâm trắc toàn diện một hoặc một số thang đo lo âu chuyên biệt cho học sinh trung học cơ sở Việt Nam — bao gồm độ tin cậy, độ hiệu lực hội tụ và phân biệt, ngưỡng cắt riêng cho dân số trong nước. Việc lựa chọn có thể từ các bộ thang quốc tế đã có nền tảng (RCADS-25, SCARED, SCAS) thông qua quy trình dịch — dịch ngược — thử nghiệm pilot — đánh giá nhân tố — xác lập ngưỡng cắt. Bộ thang chuẩn hóa này sẽ giúp các nghiên cứu trong nước có thể so sánh kết quả với nhau và với các nghiên cứu quốc tế.')

H(doc6, '2. Mở rộng phạm vi loại rối loạn lo âu được nghiên cứu', 2)
P(doc6, 'Bên cạnh dạng lo âu tổng quát đã được nghiên cứu khá rộng, các nghiên cứu tương lai cần khảo sát chuyên biệt rối loạn lo âu xã hội và rối loạn lo âu chia ly — hai dạng có ý nghĩa đặc biệt ở lứa tuổi đầu vị thành niên. Đặc biệt với các nhóm có hoàn cảnh đặc thù (học sinh dân tộc thiểu số học nội trú xa nhà, học sinh có cha mẹ đi làm xa, học sinh chuyển trường), cần có khảo sát chuyên biệt về rối loạn lo âu chia ly bằng các công cụ tâm trắc phù hợp.')

H(doc6, '3. Tích hợp đồng thời yếu tố nguy cơ và yếu tố bảo vệ trong cùng thiết kế', 2)
P(doc6, 'Phần lớn nghiên cứu Việt Nam hiện nay chỉ tập trung vào một nhóm yếu tố nguy cơ tại một thời điểm. Các nghiên cứu tương lai cần đo lường đồng thời cả năm nhóm yếu tố nguy cơ chính (áp lực học tập, nghiện điện thoại, bắt nạt thể chất, bắt nạt bằng lời và bắt nạt mạng, lòng tự trọng thấp) và năm nhóm yếu tố bảo vệ (lòng tự trọng cao, gắn bó trường học, hỗ trợ của cha mẹ, hỗ trợ từ bạn bè, biện pháp ứng phó adaptive) trong cùng một mẫu, kèm theo phân tích vai trò trung gian và điều tiết giữa các yếu tố. Phân tích chi tiết các nhóm yếu tố này được trình bày trong hai bài tổng hợp chuyên sâu của tác giả — Bài 1 về Yếu tố nguy cơ và Bài 2 về Khoảng trống can thiệp.')

H(doc6, '4. Phát triển thiết kế dọc và thử nghiệm lâm sàng ngẫu nhiên đa trung tâm', 2)
P(doc6, 'Hiện tại các nghiên cứu Việt Nam chủ yếu sử dụng thiết kế cắt ngang, không cho phép phân tích nhân quả. Các nghiên cứu ưu tiên cho giai đoạn 2026-2030 cần áp dụng thiết kế dọc với ít nhất hai đến ba thời điểm đo, cùng các thử nghiệm lâm sàng ngẫu nhiên đa trung tâm trên mẫu học sinh trung học cơ sở. Thiết kế thử nghiệm ngẫu nhiên cụm — như mô hình Bradshaw và cộng sự (2025) trên 709 học sinh tại 40 trường — đặc biệt phù hợp cho can thiệp tại trường vì cho phép phân ngẫu nhiên ở cấp trường, tránh hiệu ứng lây lan giữa học sinh cùng lớp.')

H(doc6, '5. Mở rộng mẫu sang khu vực nông thôn, miền núi và dân tộc thiểu số', 2)
P(doc6, 'Các nghiên cứu hiện có chủ yếu tập trung ở Hà Nội và Thành phố Hồ Chí Minh. Cần mở rộng mẫu sang khu vực nông thôn, miền núi, và đặc biệt nhóm học sinh dân tộc thiểu số — nơi đặc thù vùng miền có thể ảnh hưởng cả đến tỷ lệ phổ biến lẫn cơ chế tác động của các yếu tố ảnh hưởng. Nghiên cứu của Ngô Anh Vinh và cộng sự (2024) tại Lạng Sơn là một bước khởi đầu quan trọng, song cần thêm các nghiên cứu khác trên các nhóm dân tộc và vùng địa lý khác nhau.')

H(doc6, '6. Phát triển và kiểm định chương trình can thiệp bản địa hoá', 2)
P(doc6, 'Các chương trình can thiệp nhận thức — hành vi (CBT) bản địa được kiểm định trên học sinh trung học cơ sở Việt Nam còn vắng bóng. Bài 2 — Khoảng trống can thiệp đề xuất khung tích hợp ba kênh × ba tầng (lâm sàng × trường học × số; phổ quát × chọn lọc × chỉ định) cùng mô hình chăm sóc theo bậc cho Việt Nam. Các nghiên cứu tương lai cần ưu tiên: (i) thiết kế và thử nghiệm chương trình CBT trường học 8-12 buổi tích hợp vào môn Hoạt động trải nghiệm — Hướng nghiệp của Chương trình giáo dục phổ thông 2018; (ii) phát triển nền tảng CBT số tiếng Việt qua quy trình đồng thiết kế với vị thành niên, kèm cơ chế chuyển tuyến tới dịch vụ chuyên khoa khi cần; (iii) đào tạo bài bản cho đội ngũ giáo viên tâm lý học đường về sàng lọc và can thiệp sơ cấp. Luận án này, với phần đề xuất khung chương trình tập huấn phòng ngừa ở mục 3.7, đặt nền tảng ban đầu cho hướng nghiên cứu này.')

H(doc6, 'Phối hợp ba cấp độ để các kiến nghị khả thi', 2)
P(doc6, 'Để các kiến nghị trên có thể triển khai thực tế, cần phối hợp ba cấp độ. Ở cấp vĩ mô, Bộ Y tế và Bộ Giáo dục — Đào tạo phối hợp xây dựng khung chính sách sức khỏe tâm thần học đường, bao gồm tiêu chuẩn năng lực của giáo viên tâm lý và quy trình sàng lọc — chuyển tuyến. Ở cấp trung mô, các trường đại học sư phạm và viện nghiên cứu phối hợp đào tạo nhân lực chuyên môn và thực hiện các thử nghiệm can thiệp. Ở cấp vi mô, các trường trung học cơ sở là điểm tiếp xúc trực tiếp với học sinh và là nơi triển khai chương trình cụ thể. Mạng lưới nghiên cứu can thiệp liên ngành — kết nối các trường đại học, viện nghiên cứu và bệnh viện tâm thần — là điều kiện để chuẩn hóa thang đo, giao thức và quy trình báo cáo trên cả nước.')

out6 = os.path.join(OUTDIR, '06_KienNghi_HuongNghienCuu_v1_26052026.docx')
doc6.save(out6)

# ============================================================
# STATS
# ============================================================
for path in [out5, out6]:
    d = Document(path)
    w = sum(len(p.text.split()) for p in d.paragraphs)
    n_p = sum(1 for p in d.paragraphs if p.text.strip())
    n_h2 = sum(1 for p in d.paragraphs if p.style.name == 'Heading 2')
    print(f"{os.path.basename(path)}")
    print(f"  Words: {w}, Paragraphs: {n_p}, H2: {n_h2}, Size: {os.path.getsize(path)//1024}KB")
