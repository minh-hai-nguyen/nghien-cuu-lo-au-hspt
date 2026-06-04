# -*- coding: utf-8 -*-
"""Kiem tra du chu va du hinh khong - byte level + character count.
28/05/2026."""
import os, sys, io, zipfile, hashlib
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
BACKUP = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026_BEFORE_TOC.docx')
CURRENT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

print("="*75)
print("KIEM TRA BOT CHU / BOT HINH")
print("="*75)

bk = Document(BACKUP)
nw = Document(CURRENT)

# ====================================================================
# 1. CHARACTER COUNT
# ====================================================================
print("\n[1] DEM SO KY TU TUNG NGUON")

# Body paragraphs
bk_body_text = '\n'.join(p.text for p in bk.paragraphs)
nw_body_text = '\n'.join(p.text for p in nw.paragraphs)
print(f"\n  Body paragraphs:")
print(f"    Backup:  {len(bk_body_text):>10} ky tu, {len(bk.paragraphs):>5} doan")
print(f"    Current: {len(nw_body_text):>10} ky tu, {len(nw.paragraphs):>5} doan")
diff = len(nw_body_text) - len(bk_body_text)
expected_add = len('Nhấn F9 hoặc chuột phải > Update Field để cập nhật Mục lục.') + 1  # +1 for \n
print(f"    Diff:    {diff:>10} ky tu (mong doi: +{expected_add} cho TOC placeholder)")
print(f"    {'✓ KHOP' if diff == expected_add else '✗ KHONG KHOP'}")

# Table cells
def all_table_text(doc):
    parts = []
    for tb in doc.tables:
        for row in tb.rows:
            for cell in row.cells:
                parts.append(cell.text)
    return '\n'.join(parts)

bk_tbl_text = all_table_text(bk)
nw_tbl_text = all_table_text(nw)
print(f"\n  Table cells:")
print(f"    Backup:  {len(bk_tbl_text):>10} ky tu, {sum(1 for tb in bk.tables for r in tb.rows for c in r.cells):>5} cells")
print(f"    Current: {len(nw_tbl_text):>10} ky tu, {sum(1 for tb in nw.tables for r in tb.rows for c in r.cells):>5} cells")
diff_tbl = len(nw_tbl_text) - len(bk_tbl_text)
print(f"    Diff:    {diff_tbl:>10} ky tu")
print(f"    {'✓ KHONG MAT' if diff_tbl == 0 else '✗ CO MAT/THEM'}")

# Headers
def all_hf_text(doc, is_header=True):
    parts = []
    for sec in doc.sections:
        if is_header:
            hfs = [sec.header, sec.first_page_header, sec.even_page_header]
        else:
            hfs = [sec.footer, sec.first_page_footer, sec.even_page_footer]
        for hf in hfs:
            try:
                for p in hf.paragraphs:
                    parts.append(p.text)
            except: pass
    return '\n'.join(parts)

bk_h = all_hf_text(bk, True)
nw_h = all_hf_text(nw, True)
print(f"\n  Headers:")
print(f"    Backup:  {len(bk_h):>10} ky tu")
print(f"    Current: {len(nw_h):>10} ky tu")
print(f"    {'✓ KHONG MAT' if len(bk_h) == len(nw_h) else '✗ KHAC'}")

bk_f = all_hf_text(bk, False)
nw_f = all_hf_text(nw, False)
print(f"\n  Footers:")
print(f"    Backup:  {len(bk_f):>10} ky tu")
print(f"    Current: {len(nw_f):>10} ky tu")
print(f"    {'✓ KHONG MAT' if len(bk_f) == len(nw_f) else '✗ KHAC'}")

# Footnotes (read from XML)
def footnote_text(docx_path):
    with zipfile.ZipFile(docx_path) as z:
        if 'word/footnotes.xml' not in z.namelist():
            return ''
        xml = z.read('word/footnotes.xml').decode('utf-8', errors='replace')
        # Strip XML tags rough
        import re
        text = re.sub(r'<[^>]+>', '', xml)
        return text

bk_fn = footnote_text(BACKUP)
nw_fn = footnote_text(CURRENT)
print(f"\n  Footnotes (raw text):")
print(f"    Backup:  {len(bk_fn):>10} ky tu")
print(f"    Current: {len(nw_fn):>10} ky tu")
print(f"    {'✓ KHONG MAT' if bk_fn == nw_fn else '? KHAC'}")

# Total
bk_total = len(bk_body_text) + len(bk_tbl_text) + len(bk_h) + len(bk_f) + len(bk_fn)
nw_total = len(nw_body_text) + len(nw_tbl_text) + len(nw_h) + len(nw_f) + len(nw_fn)
print(f"\n  TONG (body + tables + headers + footers + footnotes):")
print(f"    Backup:  {bk_total:>10}")
print(f"    Current: {nw_total:>10}")
print(f"    Diff:    {nw_total - bk_total:>10} (mong doi: +{expected_add} cho TOC)")


# ====================================================================
# 2. WORD COUNT
# ====================================================================
print(f"\n[2] DEM TU (words)")
bk_words = len(bk_body_text.split()) + len(bk_tbl_text.split())
nw_words = len(nw_body_text.split()) + len(nw_tbl_text.split())
print(f"  Backup:  {bk_words:>8} words")
print(f"  Current: {nw_words:>8} words")
print(f"  Diff:    {nw_words - bk_words:>8} words (mong doi: +9-10 cho cau TOC placeholder)")


# ====================================================================
# 3. HINH ANH - byte-level + count
# ====================================================================
print(f"\n[3] HINH ANH (media files)")
with zipfile.ZipFile(BACKUP) as zb, zipfile.ZipFile(CURRENT) as zn:
    bk_media = sorted([n for n in zb.namelist() if 'word/media/' in n or 'word/embeddings/' in n])
    nw_media = sorted([n for n in zn.namelist() if 'word/media/' in n or 'word/embeddings/' in n])
    print(f"  Backup media count: {len(bk_media)}")
    print(f"  Current media count: {len(nw_media)}")
    missing = set(bk_media) - set(nw_media)
    added = set(nw_media) - set(bk_media)
    print(f"  Missing in current: {len(missing)} {sorted(missing)}")
    print(f"  Added in current: {len(added)} {sorted(added)}")

    # Byte hash compare for common
    common = set(bk_media) & set(nw_media)
    same = 0
    diff = 0
    diff_files = []
    for m in common:
        bk_data = zb.read(m)
        nw_data = zn.read(m)
        if bk_data == nw_data:
            same += 1
        else:
            diff += 1
            diff_files.append((m, len(bk_data), len(nw_data)))
    print(f"  Common files byte-identical: {same}/{len(common)}")
    print(f"  Modified files: {diff}")
    if diff_files:
        for f, b, n in diff_files[:5]:
            print(f"    {f}: backup={b} bytes, current={n} bytes")


# ====================================================================
# 4. DRAWING / CHART references in document.xml
# ====================================================================
print(f"\n[4] REFERENCES TO IMAGES/DRAWINGS (in document.xml)")
def count_image_refs(doc):
    n_drawing = 0
    n_blip = 0
    n_pic = 0
    for elem in doc.element.body.iter():
        tag = elem.tag
        if tag.endswith('}drawing'):
            n_drawing += 1
        elif tag.endswith('}blip'):
            n_blip += 1
        elif tag.endswith('}pic'):
            n_pic += 1
    return n_drawing, n_blip, n_pic

bd, bb, bp = count_image_refs(bk)
nd, nb, np_ = count_image_refs(nw)
print(f"  w:drawing elements: backup={bd} current={nd} {'✓' if bd==nd else '✗'}")
print(f"  a:blip elements (image references): backup={bb} current={nb} {'✓' if bb==nb else '✗'}")
print(f"  pic:pic elements (picture objects): backup={bp} current={np_} {'✓' if bp==np_ else '✗'}")


# ====================================================================
# 5. SHA256 of full text dump
# ====================================================================
print(f"\n[5] SHA256 hash cua toan bo text (body + tables)")
# Backup full text (excluding TOC placeholder which is the only addition)
bk_full = bk_body_text + bk_tbl_text
# Current full text MINUS the TOC placeholder line
nw_full = nw_body_text + nw_tbl_text
toc_line = 'Nhấn F9 hoặc chuột phải > Update Field để cập nhật Mục lục.'
nw_minus_toc = nw_full.replace('\n' + toc_line, '').replace(toc_line + '\n', '').replace(toc_line, '')
bk_hash = hashlib.sha256(bk_full.encode('utf-8')).hexdigest()
nw_hash = hashlib.sha256(nw_minus_toc.encode('utf-8')).hexdigest()
print(f"  Backup hash:           {bk_hash}")
print(f"  Current (-TOC) hash:   {nw_hash}")
if bk_hash == nw_hash:
    print(f"  ✓ HASH KHOP HOAN TOAN - khong mat chu nao")
else:
    print(f"  ✗ HASH KHAC - co chu bi mat hoac thay doi")
    # Try to find first diff
    for i in range(min(len(bk_full), len(nw_minus_toc))):
        if bk_full[i] != nw_minus_toc[i]:
            print(f"  First diff at char {i}:")
            print(f"    backup [{max(0,i-30):>}: ...{bk_full[max(0,i-30):i+30]!r}...")
            print(f"    current[{max(0,i-30):>}: ...{nw_minus_toc[max(0,i-30):i+30]!r}...")
            break

print()
print("="*75)
print("KET LUAN CUOI")
print("="*75)
