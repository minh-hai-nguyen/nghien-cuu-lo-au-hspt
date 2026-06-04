# -*- coding: utf-8 -*-
"""Bổ sung nội dung QT29-34 — tạo lại đầy đủ hơn"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *
OUT_DIR = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    import docx as dx
    d2 = dx.Document(path)
    t = '\n'.join([p.text for p in d2.paragraphs])
    print(f'  -> {name}: {len(t)} chars, {len(d2.tables)} tables')

# =====================================================================
# QT29 — BMC Psychiatry 2025 — NMA — VIẾT LẠI ĐẦY ĐỦ
# =====================================================================
def dich_QT29():
    print('QT29 CBT NMA (bổ sung)...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1186/s12888-025-07227-y', size=10)
    add_heading(doc, 'Hiệu quả các loại can thiệp khác nhau cho rối loạn lo âu ở trẻ em và thanh thiếu niên: Tổng quan hệ thống và phân tích tổng hợp mạng', 1)
    h = doc.add_paragraph(); r = h.add_run('Effectiveness of Different Types of Interventions for Anxiety Disorders in Children and Adolescents: A Systematic Review and Network Meta-Analysis'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Effectiveness of Different Types of Interventions for Anxiety Disorders in Children and Adolescents: A Systematic Review and Network Meta-Analysis'),
        ('Tác giả', 'Li LH, Wu JY, et al.'),
        ('Cơ quan', 'Capital University of Physical Education and Sports, Beijing; Shenyang Pharmaceutical University'),
        ('Tạp chí', 'BMC Psychiatry (Q1, IF ≈ 4,4)'),
        ('Xuất bản', '2025, Vol. 25, Article 809, 14 trang'),
        ('DOI', '10.1186/s12888-025-07227-y'),
        ('Loại NC', 'Tổng quan hệ thống + phân tích tổng hợp mạng Bayesian (NMA)'),
        ('Mẫu', '30 RCT, 1.711 người tham gia, 12 quốc gia, VTN tuổi TB 12,4, 62% nữ'),
    ])
    add_page_ref(doc, '1–14', 'BMC Psychiatry', '2025, 25:809')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Bối cảnh: Rối loạn lo âu phổ biến nhất ở trẻ em/VTN (4,4% ở 10–14; 5,5% ở 15–19 tuổi). Tỷ lệ tăng sau COVID. Tỷ lệ thuyên giảm sau điều trị chỉ ~50%. Phương pháp: Tìm kiếm hệ thống RCT trong PubMed, EMBASE, Cochrane, PsycINFO. NMA Bayesian so sánh 4 can thiệp: ACT, CBT, hoạt động thể chất (PE), thực tế ảo (VRET), và đối chứng. 30 RCT, 1.711 người tham gia.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Tất cả 4 can thiệp đều hiệu quả. ACT xếp hạng 1 (MD = −3,83; SUCRA = 0,69), CBT xếp hạng 2 (MD = −3,64, KTC 95%: −7,36 đến −0,48; SUCRA = 0,66), VRET xếp hạng 3 (MD = −2,53; SUCRA = 0,51), PE xếp hạng 4 (MD = −2,16; SUCRA = 0,51). CBT là can thiệp DUY NHẤT có khoảng tin cậy không bao gồm 0 — bằng chứng mạnh nhất.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['BMC Psychiatry Q1 IF ≈ 4,4. NMA Bayesian — phương pháp mạnh.','30 RCT, 1.711 VTN, 12 quốc gia.','ACT xếp hạng 1 (bất ngờ — CBT thường được coi là hàng đầu).','CBT có bằng chứng MẠNH NHẤT (KTC không qua 0).','Hoạt động thể chất (PE) cũng hiệu quả — dễ triển khai tại trường.','VRET hiệu quả — hướng tương lai.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['GRADE: bằng chứng tổng thể ở mức THẤP (low certainty).','ACT chỉ có 6 RCT — KTC rộng, bao gồm 0.','PE chỉ 2 RCT — rất ít.','Hầu hết RCT từ Mỹ, Úc, Âu — ít dữ liệu châu Á.','Dị chất đáng kể (τ = 5,87).']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Rối loạn lo âu phổ biến nhất ở trẻ em/VTN: 4,4% trẻ 10–14 tuổi và 5,5% trẻ 15–19 tuổi (WHO). Thường phát triển trước 14 tuổi, liên quan đến giảm chất lượng sống, bệnh đi kèm, bỏ học, và trong trường hợp nặng — tự tử. COVID-19 làm tăng tỷ lệ do cô lập xã hội, đóng cửa trường, giảm hoạt động thể chất.')
    add_p(doc, 'Tỷ lệ thuyên giảm sau điều trị chỉ ~50% → cần xác định can thiệp hiệu quả nhất. NMA cho phép so sánh nhiều can thiệp đồng thời, kể cả khi chưa so sánh trực tiếp.')

    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Tìm kiếm hệ thống: PubMed, EMBASE, Cochrane Library, PsycINFO. Tiêu chuẩn: RCT, trẻ em/VTN có lo âu, so sánh ACT/CBT/PE/VRET vs đối chứng. Mô hình Bayesian random effects. Xếp hạng SUCRA. Đánh giá chất lượng: Cochrane Risk of Bias + GRADE.')
    add_p(doc, '30 RCT được chọn từ 12 quốc gia (Mỹ 7, Úc 5, Đức 3, UK 2, Iran 2, Thụy Điển 2, khác 9). Tuổi TB = 12,40; 62% nữ. Can thiệp: ACT (6), CBT (16), PE (2), VRET (6). Thời gian can thiệp trung vị: 10 tuần, 1 lần/tuần, 60 phút/lần. Theo dõi trung vị: 3 tháng.')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xếp hạng hiệu quả can thiệp (NMA Bayesian)', 3)
    add_table(doc,
        ['Can thiệp', 'Số RCT', 'MD vs đối chứng', 'KTC 95% CrI', 'SUCRA', 'Xếp hạng TB'],
        [['ACT', '6', '−3,83', '−9,33 đến 1,51', '0,69', '2,25'],
         ['CBT', '16', '−3,64', '−7,36 đến −0,48*', '0,66', '2,31'],
         ['VRET', '6', '−2,53', '−8,23 đến 3,32', '0,51', '2,86'],
         ['PE', '2', '−2,16', '−9,99 đến 5,52', '0,51', '3,12'],
         ['Đối chứng', '—', 'Reference', '—', '—', '—']],
        widths=[2.5, 1.5, 3.0, 3.5, 2.0, 2.0])
    add_p(doc, '* CBT là can thiệp DUY NHẤT có KTC không bao gồm 0 — bằng chứng mạnh nhất. MD = Mean Difference. CrI = Credible Interval.', size=9, italic=True)

    add_heading(doc, 'Bảng 2. Giải thích các loại can thiệp', 3)
    add_table(doc,
        ['Can thiệp', 'Tên đầy đủ', 'Cơ chế', 'Ứng dụng cho VN'],
        [['ACT', 'Acceptance and Commitment Therapy', 'Chấp nhận cảm xúc tiêu cực + hành động theo giá trị', 'Mới, ít chuyên gia VN'],
         ['CBT', 'Cognitive-Behavioral Therapy', 'Tái cấu trúc nhận thức + phơi nhiễm dần', 'Phổ biến nhất, có thể đào tạo'],
         ['PE', 'Physical Exercise', 'Endorphin, BDNF, điều hòa HPA', 'Dễ triển khai tại trường VN'],
         ['VRET', 'Virtual Reality Exposure Therapy', 'Phơi nhiễm qua thực tế ảo', 'Cần công nghệ, đắt']],
        widths=[2.0, 4.0, 4.5, 3.5])

    add_heading(doc, '4. THẢO LUẬN', 2)
    add_p(doc, 'ACT xếp hạng 1 nhưng bằng chứng chưa chắc chắn (KTC bao gồm 0). CBT có bằng chứng mạnh nhất (16 RCT, KTC không qua 0) — vẫn là tiêu chuẩn vàng. VRET hứa hẹn nhưng cần thêm RCT ở VTN. PE đáng chú ý — dễ triển khai, ít chi phí, có cơ chế sinh học rõ ràng (endorphin, BDNF, HPA).')
    add_p(doc, 'GRADE: bằng chứng tổng thể ở mức THẤP — do thiên lệch xuất bản và imprecision (KTC rộng). Cần thêm RCT chất lượng cao, đặc biệt cho ACT, PE, VRET.')

    add_heading(doc, '5. KẾT LUẬN', 2)
    add_p(doc, 'NMA 30 RCT (1.711 VTN) cho thấy ACT, CBT, PE, VRET đều hiệu quả cho rối loạn lo âu trẻ em/VTN. ACT xếp hạng cao nhất (SUCRA 0,69), CBT có bằng chứng mạnh nhất. PE hiệu quả và dễ triển khai. Tại VN, CBT nhóm và PE tại trường là can thiệp khả thi nhất — cần RCT xác nhận.')

    add_abbreviation_table(doc, [('NMA','Network Meta-Analysis'),('ACT','Acceptance and Commitment Therapy — Liệu pháp Chấp nhận và Cam kết'),('CBT','Cognitive-Behavioral Therapy — Liệu pháp Nhận thức–Hành vi'),('PE','Physical Exercise — Hoạt động Thể chất'),('VRET','Virtual Reality Exposure Therapy — Liệu pháp Phơi nhiễm Thực tế Ảo'),('SUCRA','Surface Under the Cumulative Ranking'),('MD','Mean Difference — Khác biệt Trung bình'),('CrI','Credible Interval — Khoảng Tin cậy (Bayesian)'),('BDNF','Brain-Derived Neurotrophic Factor'),('HPA','Hypothalamic-Pituitary-Adrenal axis'),('GRADE','Grading of Recommendations Assessment, Development and Evaluation')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['NMA Bayesian — phương pháp mạnh, so sánh gián tiếp nhiều can thiệp.','30 RCT, 12 quốc gia, 1.711 VTN.','Bao gồm ACT và PE — ít NMA nào làm.','Consistency test OK (DIC consistent < inconsistent).','Funnel plot + Egger p = 0,91 — không thiên lệch xuất bản.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế chi tiết:', bold=True)
    for s in ['GRADE: bằng chứng THẤP — imprecision lớn (KTC rộng, bao gồm 0 cho ACT, PE, VRET).','PE chỉ 2 RCT — quá ít để kết luận.','Hầu hết RCT từ phương Tây (Mỹ 7, Úc 5) — ít châu Á.','Dị chất đáng kể (τ = 5,87) — các NC rất khác nhau về thiết kế.','Không phân biệt loại lo âu (GAD vs SAD vs phobia cụ thể).','Tuổi TB 12,4 — chủ yếu trẻ em, ít VTN lớn (15–17).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Khoảng trống NC / Gap:', bold=True)
    for s in ['VN CHƯA CÓ RCT nào cho can thiệp lo âu trẻ em/VTN — GAP CỰC LỚN.','CBT nhóm tại trường VN — khả thi nhất, cần RCT.','PE tại trường VN — dễ triển khai nhất (thể dục, yoga, mindfulness kết hợp vận động). Cần thử nghiệm.','ACT mới ở VN — cần đào tạo chuyên gia trước khi thử nghiệm.','So sánh hiệu quả can thiệp ở VN vs phương Tây — có thể khác do văn hóa.']:
        add_red(doc, f'• {s}')
    save(doc, '29_CBT_NetworkMeta_2025_BMCPsych.docx')

# =====================================================================
# QT30 — GBD Trends — VIẾT LẠI ĐẦY ĐỦ
# =====================================================================
def dich_QT30():
    print('QT30 GBD Trends (bổ sung)...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.04.001', size=10)
    add_heading(doc, 'Xu hướng rối loạn trầm cảm và lo âu ở VTN và thanh niên (10–24 tuổi) từ 1990 đến 2021: Phân tích Gánh nặng Bệnh tật Toàn cầu', 1)
    h = doc.add_paragraph(); r = h.add_run('Trends in Depressive and Anxiety Disorders among Adolescents and Young Adults (aged 10–24) from 1990 to 2021: A Global Burden of Disease Study Analysis'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Zhang Dongjun, Wu Mingyue, Li Xinqi, Wang Lina, Wu Jiali, Jin Mengyao'),
        ('Cơ quan', 'Đại học Y Xinxiang, Trung Quốc'),
        ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
        ('Xuất bản', '2025, 14 trang'),
        ('DOI', '10.1016/j.jad.2025.04.001'),
        ('Loại NC', 'Phân tích xu hướng sinh thái hồi cứu — Joinpoint regression, GBD 2021'),
        ('Phạm vi', '204 quốc gia, 1990–2021, VTN và thanh niên 10–24 tuổi'),
    ])
    add_page_ref(doc, '1–14', 'J Affective Disorders', '2025')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Mục tiêu: Đánh giá xu hướng dài hạn gánh nặng trầm cảm và lo âu ở 10–24 tuổi từ 1990 đến 2021, kiểm tra bất bình đẳng theo SDI.')
    p = doc.add_paragraph()
    r = p.add_run('Phương pháp: Phân tích xu hướng sinh thái hồi cứu sử dụng Joinpoint regression — ước tính AAPC cho tỷ lệ mới mắc và DALYs. Bất bình đẳng đánh giá bằng SII và CI. Phân tầng: tuổi, giới, khu vực, quốc gia, SDI.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: AAPC toàn cầu: trầm cảm 0,97% (KTC: 0,75–1,19), lo âu 0,84% (KTC: 0,48–1,21). Tăng tốc 2014–2021, đỉnh 2019. Trầm cảm: từ 3.085,5/100K (1990) lên 3.909,9/100K (2021). Lo âu: từ 708,0 lên 883,1/100K. Nữ gánh nặng cao hơn nam. Tăng nhanh nhất ở nhóm 10–14 tuổi (AAPC 1,44%). Nước SDI cao tăng nhanh nhất.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['J Affective Disorders Q1 IF ≈ 6,6.','GBD 2021 — 204 quốc gia, 31 năm, dữ liệu toàn diện nhất.','Joinpoint regression — xác định điểm thay đổi xu hướng (2014 tăng tốc).','AAPC: trầm cảm 0,97%, lo âu 0,84% — tăng liên tục.','Nhóm 10–14 tuổi tăng nhanh nhất (AAPC 1,44%) — rất đáng lo.','Phân tích bất bình đẳng (SII, CI) — ít NC nào làm.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['GBD dựa trên mô hình — nhiều nước (VN) thiếu dữ liệu gốc.','Sinh thái hồi cứu — không nhân quả.','Không tách loại lo âu (GAD, SAD, phobia).']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Trầm cảm là nguyên nhân thứ 2 gây YLDs toàn cầu, lo âu xếp thứ 6 (Lancet Psychiatry, 2024). Tuy nhiên, NC về xu hướng gánh nặng ở VTN 10–24 tuổi vẫn hạn chế, đặc biệt về mối quan hệ với SDI. Nghiên cứu nhằm: (1) phân tích xu hướng theo tuổi, giới, khu vực, SDI; (2) đánh giá bất bình đẳng giữa các nước; (3) cung cấp bằng chứng cho chính sách.')

    add_heading(doc, '2. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xu hướng toàn cầu trầm cảm và lo âu ở 10–24 tuổi', 3)
    add_table(doc,
        ['Rối loạn / Phân nhóm', 'Mới mắc 1990 (/100K)', 'Mới mắc 2021 (/100K)', 'AAPC (%)', 'p'],
        [['Trầm cảm — tổng', '3.085,5 (2.307,8–4.105,6)', '3.909,9 (2.841,3–5.171,1)', '0,97 (0,75–1,19)', '<0,001'],
         ['Trầm cảm — nam', '2.265,2', '2.991,7', '1,07', '<0,001'],
         ['Trầm cảm — nữ', '3.932,7', '4.875,1', '0,93', '<0,001'],
         ['Trầm cảm — 10–14y', '—', '—', '1,44', '<0,001'],
         ['Lo âu — tổng', '708,0 (511,5–939,3)', '883,1 (636,9–1.167,1)', '0,84 (0,48–1,21)', '<0,001'],
         ['Lo âu — 10–14y', '397,1', '547,2', '1,12', '<0,001'],
         ['Lo âu — 15–19y', '943,6', '1.180,8', '0,93', '<0,001'],
         ['Lo âu — 20–24y', '1.105,8', '1.287,1', '0,60', '<0,001']],
        widths=[3.5, 3.0, 3.0, 2.5, 1.5])

    add_heading(doc, 'Bảng 2. Xu hướng theo khu vực (AAPC trầm cảm mới mắc, cao nhất)', 3)
    add_table(doc,
        ['Khu vực', 'AAPC Trầm cảm', 'AAPC Lo âu', 'Ghi chú'],
        [['Bắc Mỹ thu nhập cao', '2,30 (1,53–3,07)', '—', 'Cao nhất trầm cảm'],
         ['Mỹ Latinh Trung', '2,00 (1,72–2,28)', '1,48 (1,18–1,79)', ''],
         ['Mỹ Latinh Andes', '1,68 (1,22–2,14)', '1,91 (1,82–2,00)', 'Cao nhất lo âu'],
         ['Mỹ Latinh nhiệt đới', '—', '1,61 (1,42–1,81)', ''],
         ['Đông Á', 'Trung bình', 'Trung bình', 'Bao gồm VN, TQ'],
         ['SDI cao', 'Cao nhất', 'Cao nhất', 'Tương quan SDI'],
         ['SDI thấp', 'Thấp nhất', 'Thấp nhất', 'Có thể do thiếu dữ liệu']],
        widths=[4.0, 3.0, 3.0, 3.5])

    add_heading(doc, '3. THẢO LUẬN VÀ KẾT LUẬN', 2)
    add_p(doc, 'Xu hướng tăng liên tục 31 năm — tăng tốc từ 2014, đỉnh 2019 (trước COVID). Nữ gánh nặng cao hơn nhưng nam tăng nhanh hơn. Nhóm 10–14 tuổi tăng nhanh nhất (AAPC 1,44%) — rất đáng lo, gợi ý rối loạn khởi phát ngày càng sớm. Bất bình đẳng: nước SDI cao tăng nhanh nhất — có thể do phát hiện + MXH + áp lực học tập.')
    add_p(doc, 'Dữ liệu GBD 204 quốc gia 31 năm xác nhận trầm cảm và lo âu ở 10–24 tuổi là vấn đề SKTT toàn cầu ngày càng nghiêm trọng, cần ưu tiên can thiệp đặc biệt cho nhóm trẻ nhất (10–14) và nữ.')

    add_abbreviation_table(doc, [('GBD','Global Burden of Disease'),('AAPC','Average Annual Percent Change'),('DALYs','Disability-Adjusted Life Years'),('SDI','Socio-Demographic Index'),('YLDs','Years Lived with Disability'),('SII','Slope Index of Inequality'),('CI','Concentration Index'),('UI','Uncertainty Interval')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['GBD 2021 — toàn diện nhất, 204 quốc gia, 31 năm.','Joinpoint regression — xác định chính xác điểm thay đổi (2006, 2010, 2014, 2019).','Phân tích bất bình đẳng (SII, CI) — đóng góp quan trọng.','Phân tầng chi tiết: tuổi × giới × khu vực × SDI.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['GBD dựa trên MÔ HÌNH ước tính — nhiều nước (VN) thiếu dữ liệu khảo sát quốc gia → ước tính có thể không chính xác.','Sinh thái hồi cứu — xu hướng ở cấp quốc gia, không phải cá nhân.','Không tách loại lo âu — GAD, SAD, phobia được gộp chung.','SDI thấp tăng chậm — có thể do THIẾU DỮ LIỆU chứ không phải thực sự thấp.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN nằm trong SDI trung bình — ước tính GBD cho VN chưa được xác nhận bằng khảo sát quốc gia (ngoại trừ V-NAMHS 2022).','Cần khảo sát quốc gia lặp lại ở VN để xác nhận xu hướng GBD.','Nhóm 10–14 tuổi tăng nhanh nhất — VN cần NC riêng cho THCS (chưa có nhiều dữ liệu).']:
        add_red(doc, f'• {s}')
    save(doc, '30_GBD_Trends_10-24y_2025.docx')

# RUN
dich_QT29()
dich_QT30()
print('Done QT29-30 (bổ sung)!')
