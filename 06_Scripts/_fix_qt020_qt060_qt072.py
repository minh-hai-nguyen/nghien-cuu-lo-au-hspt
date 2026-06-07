"""Create FIXED versions of QT020, QT060, QT072 summaries with facts verified against PDF."""
import sys, io
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_LINE_SPACING

OUT_DIR = r"c:/Users/OS/OneDrive/read_books/Lo-au/Tom-tat-tung-bai"

FOOTER_NOTE = "Đã sửa lỗi factual (07/06/2026) — đối chiếu trực tiếp PDF gốc"


def new_doc():
    d = Document()
    # Normal style: Times New Roman 12pt, line spacing 1.5
    st = d.styles['Normal']
    st.font.name = 'Times New Roman'
    st.font.size = Pt(12)
    st.paragraph_format.line_spacing = 1.5
    # Strip metadata
    cp = d.core_properties
    cp.author = ''
    cp.title = ''
    cp.keywords = ''
    cp.subject = ''
    cp.comments = ''
    cp.last_modified_by = ''
    cp.category = ''
    return d


def add_para(doc, text, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(text)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(12)
    r.bold = bold
    return p


def add_heading(doc, text):
    add_para(doc, text, bold=True)


def add_footer(doc, note):
    sec = doc.sections[0]
    f = sec.footer
    # Use first paragraph if exists
    if f.paragraphs:
        p = f.paragraphs[0]
        p.text = ''
    else:
        p = f.add_paragraph()
    p.paragraph_format.line_spacing = 1.5
    r = p.add_run(note)
    r.font.name = 'Times New Roman'
    r.font.size = Pt(10)
    r.italic = True


# ============ QT020 ============
d = new_doc()
add_para(d, "QT020 — Tóm tắt bài nghiên cứu (đã đối chiếu PDF gốc 07/06/2026)", bold=True)

add_heading(d, "Tên công trình")
add_para(d, '"CBT treatment delivery formats for generalized anxiety disorder: a systematic review and network meta-analysis of randomized controlled trials" — Liu S, Xiao H, Duan Y, Shi L, Wang P, Cao L, Li H, Huang X, Qiu C (2025). Translational Psychiatry 15:197. DOI: 10.1038/s41398-025-03414-3.')

add_heading(d, "Phương pháp nghiên cứu")
add_para(d, "• Systematic review + network meta-analysis (NMA) — đăng ký PROSPERO CRD42023493949")
add_para(d, "• Tìm kiếm: MEDLINE, Embase, PsycINFO, Web of Science đến 09/2023")
add_para(d, "• Mẫu: 52 RCT, N = 4.361 bệnh nhân GAD (5 cụm phương thức: cá nhân, nhóm, từ xa, TAU, waitlist)")
add_para(d, "• Tuổi trung bình: Abstract ghi 43 tuổi; Body (trang 2) ghi 40,8 tuổi — đối chiếu PDF có chênh do tính khác arms; ưu tiên trích body 40,8 tuổi khi vào luận án")
add_para(d, "• Tỷ lệ nữ: Abstract 69,7 %; Body 67,1 % — tương tự, ưu tiên body")
add_para(d, "• Công cụ phân tích: random-effects, pairwise + NMA (Stata mvmeta), SUCRA, ROB 2")
add_para(d, "• Outcome efficacy = SMD trên thang lo âu (continuous); acceptability = RR all-cause discontinuation (dichotomous)")

add_heading(d, "Phát hiện chính (efficacy outcome — thang lo âu)")
add_para(d, "• Individual CBT > Remote CBT: SMD 0,96 [95% CI 0,13; 1,79] — abstract của PDF ghi rõ thuộc về efficacy (\"in relieving anxiety symptoms of GAD\"). Lưu ý: body trang 3 có cụm \"Individual CBT was more acceptable than remote CBT\" dùng cùng giá trị 0,96 — đây nhiều khả năng là LỖI BIÊN TẬP trong PDF (acceptability không thể dùng SMD vì là RR), nên đọc theo abstract.")
add_para(d, "• Individual CBT > Treatment-as-usual (TAU): SMD 1,12 [0,24; 2,00]; SUCRA 82,8 %")
add_para(d, "• Individual CBT > Waitlist: SMD 1,62 [1,03; 2,22]")
add_para(d, "• Group CBT > Waitlist: SMD 1,65 [0,47; 2,84]; SUCRA 86,9 %")
add_para(d, "• Remote CBT KHÔNG vượt trội TAU (SMD 0,16 [-1,04; 1,36]) hay Waitlist (SMD 1,16 [-0,64; 2,95]) — CI rộng, hiệu lực không kết luận")
add_para(d, "• 75 % RCT có low/unclear risk of bias (39/52 trial)")

add_heading(d, "Phát hiện chính (acceptability outcome — dropout)")
add_para(d, "• KHÔNG có khác biệt có ý nghĩa thống kê giữa các format CBT về tỷ lệ bỏ dở trial (page 1 abstract + page 4 body)")

add_heading(d, "Điểm mạnh")
add_para(d, "• NMA 52 RCT — cỡ bằng chứng lớn cho GAD người lớn")
add_para(d, "• Translational Psychiatry — Q1 Nature group")
add_para(d, "• So sánh 5 nodes: individual, group, remote, TAU, waitlist")
add_para(d, "• Effect sizes có CI rõ ràng, SUCRA, side-split test cho consistency")

add_heading(d, "Hạn chế")
add_para(d, "⚠ SCOPE: đối tượng là NGƯỜI LỚN GAD (tuổi TB ~41) — KHÔNG PHẢI vị thành niên (VTN). Body PDF nêu chỉ 2/52 trial bao gồm trẻ ≤18 và 5/52 bao gồm người ≥65. Áp dụng cho HS Việt Nam cần ngoại suy thận trọng.")
add_para(d, "• Tỷ lệ nữ 67-70 % — mẫu nghiêng nữ")
add_para(d, "• 25 % RCT high risk of bias (sensitivity analysis loại bỏ 5 trial vẫn ổn định)")
add_para(d, "• Group CBT chỉ có 7 arms (343 BN) — bằng chứng mỏng cho group")
add_para(d, "• Tìm kiếm đến 09/2023 — RCT mới hơn chưa được gộp")

add_heading(d, "Phản biện ngắn")
add_para(d, "NMA lớn về CBT FORMAT — nhưng đối tượng là NGƯỜI LỚN GAD. Kết quả \"Individual CBT > Remote/TAU/Waitlist; Group CBT > Waitlist\" hữu ích làm REFERENCE HIERARCHY về thứ tự ưu tiên format, nhưng effect size (SMD 0,96-1,65) KHÔNG THỂ áp dụng trực tiếp cho HS/VTN vì: (1) tâm lý phát triển khác; (2) động lực trị liệu khác (VTN ít tự giác hơn); (3) Individual CBT với VTN cần involve phụ huynh. Cần NMA tương tự nhưng restrict to <18 tuổi (so sánh với QT029 Li 2025).")

add_heading(d, "Hướng nghiên cứu tiếp theo")
add_para(d, "Lặp lại NMA chỉ với RCT VTN/trẻ em. Phân tích subgroup format theo tuổi. Đánh giá Remote CBT trong bối cảnh hậu COVID khi telehealth phổ biến hơn.")

add_heading(d, "Đánh giá chất lượng: ⭐⭐⭐ (3/5)")

add_heading(d, "Key fact cho RAG")
add_para(d, "GAD người lớn (~41 tuổi). Individual CBT > Group ≈ Remote > TAU > Waitlist xét theo efficacy SMD. Acceptability không khác biệt giữa các format. KHÔNG phải VTN — dùng tham chiếu hierarchy, KHÔNG copy effect size cho luận án CTH.")

add_heading(d, "Tham khảo")
add_para(d, "• Liu S, et al. (2025). CBT treatment delivery formats for generalized anxiety disorder: a systematic review and network meta-analysis of randomized controlled trials. Transl Psychiatry 15:197. https://doi.org/10.1038/s41398-025-03414-3")
add_para(d, "• PROSPERO CRD42023493949")
add_para(d, "• PDF gốc: 02_Papers-goc/The-gioi_Khac/QT020_Liu_CBT_Delivery_GAD_NMA_2025.pdf (trang 1-4 abstract+kết quả)")

add_footer(d, FOOTER_NOTE)
d.save(f"{OUT_DIR}/QT020_Liu_CBT_Delivery_GAD_NMA_2025_FIXED_07062026.docx")
print("Saved QT020 FIXED")


# ============ QT060 ============
d = new_doc()
add_para(d, "QT060 — Tóm tắt bài nghiên cứu (đã đối chiếu PDF gốc 07/06/2026)", bold=True)

add_heading(d, "Tên công trình")
add_para(d, '"Rising global burden of anxiety disorders among adolescents and young adults: trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021" — Bie F, Yan X, Xing J, Wang L, Xu Y, Wang G, Wang Q, Guo J, Qiao J, Rao Z (2024). Frontiers in Psychiatry 15:1489427. DOI: 10.3389/fpsyt.2024.1489427.')

add_heading(d, "Phương pháp")
add_para(d, "• Phân tích thứ cấp GBD 2021 (Global Health Data Exchange / GHDx)")
add_para(d, "• Phạm vi: 204 quốc gia & vùng lãnh thổ, 21 region, 5 nhóm SDI (low, low-middle, middle, high-middle, high)")
add_para(d, "• Đối tượng: 10-24 tuổi (3 sub-group: 10-14, 15-19, 20-24)")
add_para(d, "• Outcome: incidence, prevalence, DALYs — rate / 100.000 + counts")
add_para(d, "• Công cụ: Joinpoint regression (APC, AAPC, EAPC, 95% CI); SDI analysis; báo cáo theo STROBE")
add_para(d, "• Tổng dân số 10-24 tuổi trong phân tích: 399.121.466 (nam 40,99 %; nữ 59,01 %)")

add_heading(d, "Phát hiện chính — TOÀN CẦU")
add_para(d, "• Số ca mới (incident cases) lo âu 10-24 tuổi 2021: 16.670.879 (95% UI 12,02-22,03 triệu)")
add_para(d, "• Số ca mới TĂNG 52,00 % (1990-2021), 95% UI 49-56 %")
add_para(d, "• Tỷ lệ mới mắc: 708,02 → 883,10 / 100.000 (EAPC 0,20; 95% CI 0,01-0,39)")
add_para(d, "• Tăng nhiều nhất ở nhóm 20-24 tuổi: +28,33 %; tăng ít nhất 10-14 tuổi: +21,58 %")
add_para(d, "• Nhóm 10-14 tuổi DUY TRÌ tỷ lệ cao nhất xuyên suốt (760,96 năm 1990; 925,07 năm 2021)")
add_para(d, "• Nữ > nam ở mọi nhóm tuổi (consistent)")
add_para(d, "• Prevalence toàn cầu 10-24 tuổi: 4.120,60 → 4.976,61 / 100.000 (EAPC 0,16)")
add_para(d, "• DALYs tăng 104,85 % (505,59 → 610,44 / 100.000); tăng lớn nhất nhóm 20-24 tuổi")
add_para(d, "• COVID-19: tăng tốc rõ rệt giai đoạn 2019-2021 sau giai đoạn tăng chậm")

add_heading(d, "Phát hiện chính — THEO SDI (lưu ý: cases vs rates)")
add_para(d, "• Middle SDI region có SỐ CA tuyệt đối cao nhất 2021 (incidence 5.018.673 cases; prevalence 29.018.755 cases) — vì middle SDI có dân số 10-24 lớn")
add_para(d, "• High SDI region có TỶ LỆ MỚI MẮC cao nhất (1.225,83/100.000) và TĂNG NHIỀU NHẤT (EAPC incidence 0,465; EAPC prevalence 0,45)")
add_para(d, "• High-middle SDI có số ca THẤP nhất (2.084.464)")

add_heading(d, "Phát hiện chính — THEO 21 REGION GBD")
add_para(d, "• Nam Á (South Asia) có SỐ CA TUYỆT ĐỐI cao nhất 2021: 3.478.267 incidence; 16.949.401 prevalence (chủ yếu Ấn Độ)")
add_para(d, "• Tropical Latin America có TỶ LỆ MỚI MẮC cao nhất: 1.586,81/100.000; tăng nhanh nhất (EAPC incidence 0,836)")
add_para(d, "• Western Europe có TỶ LỆ PREVALENCE cao nhất (9.368,64/100.000)")
add_para(d, "• Central Latin America có EAPC prevalence cao nhất (0,68); Central Asia có tỷ lệ thấp nhất")
add_para(d, "• Đông Á (East Asia) là region DUY NHẤT có incidence GIẢM (EAPC -0,432; 95% CI -0,657 đến -0,208) và prevalence cũng GIẢM (EAPC -0,29)")
add_para(d, "• Ấn Độ: số ca cao nhất quốc gia; Mexico: mức tăng lớn nhất (theo abstract)")

add_heading(d, "Yếu tố nguy cơ")
add_para(d, "• Bullying victimization được nêu là yếu tố nguy cơ chính, đặc biệt ở các region có gánh nặng lo âu cao (abstract + RQ3)")

add_heading(d, "Điểm mạnh")
add_para(d, "• GBD 2021 — dữ liệu cập nhật mới nhất (released 2024)")
add_para(d, "• Phân tầng đa chiều: tuổi, giới, region, country, SDI, COVID")
add_para(d, "• 32 năm theo dõi (1990-2021)")
add_para(d, "• Joinpoint regression cho phép phát hiện điểm gãy xu hướng")
add_para(d, "• Frontiers in Psychiatry Q1 Open Access; STROBE-compliant")

add_heading(d, "Hạn chế")
add_para(d, "• GBD dùng mô hình DisMod-MR ước lượng — không phải đo trực tiếp; tin cậy thấp ở LMIC có dữ liệu sơ sài (Việt Nam thuộc nhóm này)")
add_para(d, "• Anxiety disorders trong GBD là tổng hợp (DSM-IV-TR 300.0-300.3, F40-42, F43.0-1, F93.0-93.2, F93.8) — không tách subtype (GAD, SAD, phobia)")
add_para(d, "• Bullying được nêu là risk factor chính nhưng phân tích chi tiết hạn chế trong abstract")
add_para(d, "• Cross-sectional thiết kế — không cho phép suy luận nhân quả")

add_heading(d, "Phản biện ngắn")
add_para(d, "Bài QUAN TRỌNG để có BASELINE GLOBAL trends 10-24 tuổi 1990-2021. Xu hướng +52 % toàn cầu, nhưng tách theo region xuất hiện 2 chiều ngược: ĐÔNG Á GIẢM (EAPC -0,43, chủ yếu Trung Quốc + Hàn Quốc + Nhật) trong khi Tropical Latin America tăng mạnh nhất (EAPC +0,84). Việt Nam thuộc South-East Asia GBD region (KHÔNG phải East Asia), nằm trong khu vực ASEAN có ít dữ liệu trực tiếp. Đề xuất dùng GBD trends làm BACKDROP cho luận điểm: \"Việt Nam cần khảo sát quốc gia lặp lại để có xu hướng riêng, vì hiện tại GBD ước lượng cho VN dựa modelling, không phải khảo sát local.\"")

add_heading(d, "Hướng nghiên cứu tiếp theo")
add_para(d, "So sánh trends VN với Đông Á (giảm) và Tropical Latin America (tăng mạnh) để xác định mô hình VN. Phân tích bullying VN qua các NC hiện có (VN015 nếu có; VN022 UNICEF). Cập nhật khi GBD 2023/2024 ra mắt.")

add_heading(d, "Đánh giá chất lượng: ⭐⭐⭐⭐⭐ (5/5)")

add_heading(d, "Key fact cho RAG")
add_para(d, "GBD 2021 anxiety 10-24y: +52 % global 1990-2021 (708→883/100k incident rate); 16,67 triệu ca mới 2021. Nữ > nam. Middle SDI cao nhất về CASES (do dân số lớn); High SDI cao nhất về RATE và tăng mạnh nhất. Tropical Latin America cao nhất + tăng nhanh nhất theo region; Đông Á là region duy nhất GIẢM (EAPC -0,43). Nam Á dẫn đầu cases tuyệt đối (3,48 triệu, chủ yếu Ấn Độ). Bullying là yếu tố nguy cơ chính.")

add_heading(d, "Tham khảo")
add_para(d, "• Bie F, Yan X, Xing J, Wang L, Xu Y, Wang G, Wang Q, Guo J, Qiao J, Rao Z (2024). Rising global burden of anxiety disorders among adolescents and young adults: trends, risk factors, and the impact of socioeconomic disparities and COVID-19 from 1990 to 2021. Front Psychiatry 15:1489427. https://doi.org/10.3389/fpsyt.2024.1489427")
add_para(d, "• PDF gốc: 02_Papers-goc/The-gioi_Au-My-Uc/QT060_Bie_GlobalAnxiety_GBD_10-24y_1990_2021_FrontPsych_2024.pdf (Bảng 1 trang 5 — incidence theo SDI và 21 region)")

add_footer(d, FOOTER_NOTE)
d.save(f"{OUT_DIR}/QT060_Bie_GlobalAnxiety_GBD_10-24y_1990_2021_FrontPsych_2024_FIXED_07062026.docx")
print("Saved QT060 FIXED")


# ============ QT072 ============
d = new_doc()
add_para(d, "QT072 — Tóm tắt bài nghiên cứu (đã đối chiếu PDF gốc 07/06/2026)", bold=True)

add_heading(d, "Tên công trình")
add_para(d, '"Cyberbullying Victimization and Mental Health Symptoms Among Children and Adolescents: A Meta-Analysis of Longitudinal Studies" — Jungup Lee, Hyekyung Choo, Yijing Zhang, Hoi Shan Cheung, Qiyang Zhang, Rebecca P. Ang (2025). Trauma, Violence, & Abuse 27(2):391-406. DOI: 10.1177/15248380241313051. Q1 SAGE.')
add_para(d, "Ghi chú: tóm tắt v1 trước đó CHỈ liệt kê 4/6 tác giả (Lee, Cheung, Zhang Q, Ang) — đã sửa bổ sung Hyekyung Choo và Yijing Zhang theo PDF trang 1.")

add_heading(d, "Phương pháp")
add_para(d, "• META-ANALYSIS DỌC ĐẦU TIÊN về cyberbullying victimization → mental health (các meta-analysis trước chỉ cắt ngang hoặc gộp với traditional bullying)")
add_para(d, "• Tìm kiếm: 8 database (PsycINFO, SCOPUS, WoS, PubMed, MEDLINE, ERIC, ProQuest, Google Scholar) + reference list, bài tiếng Anh 01/2010-06/2021")
add_para(d, "• Tiêu chí: thiết kế dọc ≥2 wave; trẻ 8-19 tuổi; cyberbullying victimization là IV; mental health là DV (depression, anxiety, distress)")
add_para(d, "• Số nghiên cứu: 27 longitudinal studies")
add_para(d, "• Cỡ mẫu: PDF abstract ghi \"13.497 children and adolescents\"; PDF body trang 6 + Bảng 2 ghi \"total sample size = 27.133 participants\" (overall mental health). Hai con số KHÔNG khớp trong cùng bài báo — đây là LỖI BIÊN TẬP của PDF; trước đây tóm tắt v1 chỉ ghi 27.133. Khi trích vào luận án nên TRÍCH CẢ HAI và ghi chú \"abstract vs body discrepancy\".")
add_para(d, "• Quốc gia: 9 nước — Mỹ (n=12 studies, 44,4 %), Trung Quốc (7, 25,9 %), Tây Ban Nha (2), Thụy Điển (1), Hà Lan (1), Bỉ (1), Thụy Sĩ (1), Úc (1), New Zealand (1)")
add_para(d, "• Phân tích: random-effects model; Q-test heterogeneity; subgroup theo culture (Western vs Asian); meta-regression cho age, female %, publication year, time interval; Egger's test + Rosenberg's Fail-Safe N cho publication bias")
add_para(d, "• Tài trợ: MOE Singapore AcRF Tier 1 (FY2019-FRC2-002)")

add_heading(d, "Kết quả CỐT LÕI (Bảng 2 PDF)")
add_para(d, "• Cyberbullying → Mental health overall: r = 0,232 [95% CI 0,190-0,274]; k=27; N=27.133; I² = 92,25 % (heterogeneity rất cao); Q = 335,48; p < 0,001")
add_para(d, "• Depression: r = 0,269 [0,211-0,325]; k=23; N=21.481; I² = 94,91 %")
add_para(d, "• ANXIETY: r = 0,229 [0,162-0,294]; k=16; N=13.987; I² = 94,07 %")
add_para(d, "• Loneliness: r = 0,170 [0,109-0,229]; k=4; N=2.780; I² = 62,00 %")
add_para(d, "• Body image / negative cognition: r = 0,024 [-0,173; 0,219]; k=3; N=5.838 (KHÔNG có ý nghĩa)")
add_para(d, "• Somatic complaints / sleep / stress: r = 0,233 [0,202-0,264]; k=3; N=5.784")
add_para(d, "• KHÔNG có publication bias")

add_heading(d, "Phân tích moderator (meta-regression Bảng 3)")
add_para(d, "• Tuổi: β = 0,037 (p<0,001) — tuổi càng cao tác động càng mạnh")
add_para(d, "• Tỷ lệ nữ: β = -0,003 (p=0,006) — tỷ lệ nữ cao hơn → tác động giảm (NGƯỢC giả thuyết); nhóm nhiều nam thì hiệu ứng mạnh hơn")
add_para(d, "• Năm công bố: β = 0,009 (p<0,001) — bài mới hơn có effect size lớn hơn (cyberbullying ngày càng nghiêm trọng)")
add_para(d, "• Time interval: β = -0,001 (p=0,526) — KHÔNG có ý nghĩa")
add_para(d, "• Văn hóa Tây vs Á: subgroup analysis KHÔNG khác biệt (p=0,730); r Asian = 0,25 (n=7), r Western = 0,23 (n=20)")

add_heading(d, "Phát hiện ĐẶC BIỆT")
add_para(d, "• Cyberbullying là vấn đề TOÀN CẦU vượt văn hóa (Western vs Asian KHÔNG khác biệt). Internet có vẻ \"san phẳng\" khác biệt văn hóa.")
add_para(d, "• Hiệu ứng MẠNH HƠN ở tuổi cao hơn (so với phát hiện Pichel 2021 về đỉnh 14-15 tuổi)")
add_para(d, "• Hiệu ứng tăng theo năm — phản ánh cyberbullying ngày càng phổ biến và/hoặc khó tránh trong giai đoạn COVID")
add_para(d, "• Time interval KHÔNG moderator → tác động cyberbullying dai dẳng theo thời gian, không phai mờ tự nhiên")

add_heading(d, "Đánh giá chất lượng")
add_para(d, "⭐⭐⭐⭐⭐ Rất cao. Meta-analysis DỌC ĐẦU TIÊN (chỉ longitudinal, KHÔNG cắt ngang). PRISMA. Trauma, Violence, & Abuse Q1 SAGE. Đáng tin cậy cao. Lưu ý: I² > 92 % nên thận trọng khi diễn giải pooled r.")

add_heading(d, "Ý nghĩa cho VN")
add_para(d, "Bổ sung cho Moore 2017 (QT070) về traditional bullying. Cyberbullying r = 0,229 với anxiety là hiệu ứng NHỎ-VỪA (theo Cohen) NHỎ HƠN traditional bullying (Moore OR = 1,77 ≈ r ~ 0,3). Cyberbullying tác động dai dẳng (time interval không moderator). Tham chiếu cho phản biện luận án CTH: luận án chỉ đo \"bắt nạt thể chất\" (Bảng 3.21), CHƯA đo cyberbullying riêng — đây là khoảng trống cần lưu ý vì cyberbullying đang là vấn đề toàn cầu vượt văn hóa.")

add_heading(d, "Tham khảo")
add_para(d, "• Lee J, Choo H, Zhang Y, Cheung HS, Zhang Q, Ang RP (2025). Cyberbullying Victimization and Mental Health Symptoms Among Children and Adolescents: A Meta-Analysis of Longitudinal Studies. Trauma Violence Abuse 27(2):391-406. https://doi.org/10.1177/15248380241313051")
add_para(d, "• PDF gốc: 02_Papers-goc/QT072_Lee_2025_CyberbullyingMeta_TVA.pdf (trang 1 tác giả + abstract; trang 6 Statistical Analyses + Study Characteristics; trang 8 Table 2 + Table 3)")

add_footer(d, FOOTER_NOTE)
d.save(f"{OUT_DIR}/QT072_Lee_2025_CyberbullyingMeta_TVA_FIXED_07062026.docx")
print("Saved QT072 FIXED")
print("DONE")
