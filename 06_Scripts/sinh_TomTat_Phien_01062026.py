# -*- coding: utf-8 -*-
"""Sinh doc tom tat toan bo cong viec da lam trong cac phien 29/05 - 01/06/2026.
Bao gom: ra soat kho du lieu, ket qua, files quan trong, status Q1+Q3,
4 BLOCKING questions, va titles tham khao.
01/06/2026."""
import os, sys, io
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
OUT = os.path.join(ROOT, 'bai-bao-Q1', 'TomTat_Phien_01062026.docx')

d = Document()
for sec in d.sections:
    sec.top_margin = Cm(2.0); sec.bottom_margin = Cm(2.0)
    sec.left_margin = Cm(2.0); sec.right_margin = Cm(2.0)
s = d.styles['Normal']
s.font.name = 'Times New Roman'; s.font.size = Pt(11)
s.paragraph_format.line_spacing = 1.3


def H1(text):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18); p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

def H2(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(14); p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(13); r.bold = True
    r.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

def H3(text):
    p = d.add_paragraph()
    p.paragraph_format.space_before = Pt(8); p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(12); r.bold = True

def P(text, italic=False, indent=False):
    p = d.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    if indent: p.paragraph_format.first_line_indent = Cm(0.5)
    r = p.add_run(text); r.font.name = 'Times New Roman'; r.font.size = Pt(11); r.italic = italic

def B(text, level=0):
    p = d.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5 + level * 0.4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run('▸ ' + text); r.font.name = 'Times New Roman'; r.font.size = Pt(11)

def set_col_widths(table, widths_cm):
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)

def make_table(headers, rows, col_widths_cm, style='Light Grid Accent 1'):
    t = d.add_table(rows=1, cols=len(headers)); t.style = style; t.autofit = False
    hdr = t.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = h
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.font.bold = True; r.font.size = Pt(10)
    for row_data in rows:
        row = t.add_row().cells
        for i, txt in enumerate(row_data):
            row[i].text = str(txt)
            for p in row[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)
    set_col_widths(t, col_widths_cm)
    return t


# ============================================================
H1('TÓM TẮT CÔNG VIỆC ĐÃ HOÀN THÀNH')
P('Tóm tắt 4 phiên rà soát kho dữ liệu + xây dựng bố cục Q1/Q3 + chuẩn '
  'hoá thuật ngữ + tổng hợp titles tham khảo.', italic=True)
P('Giai đoạn: 29/05/2026 – 01/06/2026', italic=True)
P('Nhóm tác giả: Hang Thi Cong (1st), Nguyen Minh Duc, Duc Minh Dao '
  '(corresponding)', italic=True)


# ============================================================
H1('1. RÀ SOÁT TOÀN BỘ KHO DỮ LIỆU')

H2('1.1 Phạm vi đã rà soát')
make_table(
    ['Loại tài liệu', 'Số file', 'Coverage'],
    [
        ('PDF gốc trong 02_Papers-goc/', '147/183', '80% (36 cloud-only)'),
        ('Tóm tắt QT (Tom-tat-tung-bai/)', '73', '100%'),
        ('Bản dịch trong 03_Ban-dich/', '93 (active) + 78 metadata cleaned',
         '100%'),
        ('LA + tài liệu phái sinh (Luận án TS/)', '~20 file', '100%'),
        ('Outline Q1/Q3 (bai-bao-Q1/)', '~15 file', '100%'),
        ('CSDL (04_Co-so-du-lieu/)', '4 file', '100%'),
        ('RAG vector DB', '1 file SQLite (1.6 MB)', 'Đã inventory'),
        ('Báo cáo trong 01_Bao-cao/', '~100 file', 'Đã update thuật ngữ'),
    ],
    [5.0, 5.5, 6.5]
)

H2('1.2 Lỗi đã phát hiện và xử lý')
make_table(
    ['Loại lỗi', 'Số lượng', 'Xử lý'],
    [
        ('Số liệu thực sự sai trong QT', '1 (QT011: 44,2% → 48,7%)', 'Đã fix'),
        ('PDF mislabel filename ≠ content', '1 (QT002_Saikia → Bhardwaj)',
         'Đã đổi tên file'),
        ('Thuật ngữ "lan tỏa" cũ cho GAD', '~1.696 chỗ trong 160 file',
         'Đã batch update → "tổng quát"'),
        ('Watermark + metadata', 'Hàng trăm file',
         'Đã clean batch'),
        ('Thang đo bịa (MPAS, MPVS)', 'Phát hiện sớm các phiên trước',
         'Đã sửa thành SAS-SV + OBVQ'),
    ],
    [5.5, 5.0, 6.5]
)

H2('1.3 Chất lượng tổng thể')
B('Tỷ lệ lỗi thực tế trong QT summaries: ~1.4% (1/73 confirmed)')
B('Sample manual verify 6 case suspicious: 100% là false positive '
  '(do script chưa xử lý PDF table format hoặc derived calculations)')
B('Tất cả 17 citations chính trong Q1+Q3 đều có PDF verified trong kho')
B('Saikia 2023 PDF thật được tìm thấy (file 11_Saikia_2023_IJCM.pdf), '
  'không bị mất như báo cáo ban đầu')


# ============================================================
H1('2. CHUẨN HOÁ THUẬT NGỮ GAD: "LAN TỎA" → "TỔNG QUÁT"')

P('Theo gợi ý của Thầy Đào Minh Đức, thuật ngữ "rối loạn lo âu lan tỏa" '
  'đã được thay bằng "rối loạn lo âu tổng quát" trên toàn bộ tài liệu.',
  italic=False)

H2('2.1 Lý do chuyển đổi')
B('Bám sát ngữ nghĩa DSM-5: "Generalized" = không đặc thù theo trigger, '
  'không phải "pervasive"')
B('Tránh trùng với "Pervasive Developmental Disorder (PDD)" trong DSM-IV '
  '— đã dịch là "rối loạn phát triển lan tỏa"')
B('Logic taxonomy: GAD là dạng "general" trong nhóm rối loạn lo âu, từ đó '
  'phân tách ra lo âu chia ly, lo âu xã hội')
B('Tính phân biệt: mọi RLLA đều "lan tỏa" ảnh hưởng tới phát triển — '
  '"lan tỏa" không phải tính chất riêng của GAD')

H2('2.2 Phạm vi đã update')
make_table(
    ['Phase', 'Files', 'Replacements'],
    [
        ('Phase 1 — Core (LA + Tóm tắt + Trích yếu + Q1/Q3 outlines)',
         '15', '169'),
        ('Phase 2 — Mở rộng (báo cáo + bản dịch + working drafts)',
         '139', '1.522'),
        ('Phase 3 — Edge cases (5 file residual)', '5', '3'),
        ('CSDL master DATABASE_BAI_BAO_LO_AU.md', '1', '2'),
        ('Tổng', '160 file', '1.696 replacements'),
    ],
    [9.0, 3.0, 4.0]
)

H2('2.3 Viết tắt')
B('"RLLALT" (rối loạn lo âu lan tỏa) → "RLLATQ" (rối loạn lo âu tổng quát)')
B('Files bai-bao-khgdvn/ + backups + archive: KHÔNG đụng (an toàn theo '
  'instruction của Thầy)')


# ============================================================
H1('3. BỐ CỤC 2 BÀI Q1 + Q3 — PHƯƠNG ÁN A')

H2('3.1 Bài Q1 — BMC Psychiatry (IF 4.4, acceptance ~30%)')

make_table(
    ['Mục', 'Nội dung'],
    [
        ('Tên (tentative)',
         'Integrated risk-protective structural equation model of anxiety '
         'disorder subtypes among Vietnamese lower secondary school students: '
         'A mixed-methods study'),
        ('Tác giả',
         'Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹†'),
        ('Số từ dự kiến', '6.000–8.000'),
        ('Phương pháp', 'SEM tích hợp + Multi-group invariance + Mixed-methods'),
        ('3 đóng góp mới',
         '(1) Mô hình SEM tích hợp đầu tiên ở VN cho 7 dự báo → 3 phân loại RLLA; '
         '(2) Tính bất biến giới của lo âu chia ly (F=0,246; p=0,620); '
         '(3) Tích hợp định tính – định lượng'),
        ('Backup journal', 'BMC Psychology (IF 2.6)'),
    ],
    [4.0, 13.0]
)


H2('3.2 Bài Q3 — PLOS ONE (IF 3.7, acceptance ~50%)')

make_table(
    ['Mục', 'Nội dung'],
    [
        ('Tên (tentative)',
         'Manifestations and patterns of anxiety disorder subtypes among '
         'Vietnamese lower secondary school students: A descriptive cross-'
         'sectional study'),
        ('Tác giả',
         'Hang Thi Cong¹*, Nguyen Minh Duc², Duc Minh Dao¹†'),
        ('Số từ dự kiến', '3.500–5.000'),
        ('Phương pháp', 'Descriptive + ANOVA + Bonferroni + Cohen d/η²'),
        ('3 đóng góp mới',
         '(1) Phân tích mức độ mục đầu tiên cho RCADS phiên bản tiếng Việt; '
         '(2) Quỹ đạo phát triển theo khối lớp (SAD giảm đơn điệu 32,13 → '
         '19,46); '
         '(3) Xác định mục tiêu sàng lọc'),
        ('Backup journal', 'BMC Pediatrics (IF 2.5)'),
    ],
    [4.0, 13.0]
)


H2('3.3 Phân chia dữ liệu chống trùng lặp')

P('Theo phương án A, dữ liệu được chia theo PHƯƠNG PHÁP và CÂU HỎI nghiên '
  'cứu để tránh self-plagiarism:', italic=True)

make_table(
    ['Dữ liệu', 'Q1', 'Q3'],
    [
        ('Bảng item-level GAD/SAD/SocAD', 'Tóm tắt summary',
         'CHỦ LỰC (full table)'),
        ('Bảng demographic comparison (F, p)', 'Background',
         'Primary analysis'),
        ('Item-level yếu tố nguy cơ + bảo vệ', 'Aggregate → SEM predictors',
         'KHÔNG dùng'),
        ('SEM β coefficients (21 paths)', 'PRIMARY FINDING', 'KHÔNG dùng'),
        ('Multi-group invariance', 'NOVEL FINDING', 'Brief mention'),
        ('Qualitative interviews', 'Mixed-methods integration', 'KHÔNG dùng'),
    ],
    [5.5, 5.5, 6.0]
)


# ============================================================
H1('4. RÀ SOÁT CHI TIẾT 19 ISSUES Q1 + Q3')

H2('4.1 Tóm tắt issues')
make_table(
    ['Tier', 'Số issues', 'Owner', 'Trạng thái'],
    [
        ('Tier 1 — BLOCKING', '4 (Q1-6, Q1-8, Q3-6, Q3-9)',
         'NCS + Thầy', 'Đang chờ quyết'),
        ('Tier 2 — HIGH (đã fix outline v3)',
         '12 (Q1-1,2,3,4,5,7,9 + Q3-1,2,3,4,5)',
         'Em', 'ĐÃ FIX'),
        ('Tier 3 — MEDIUM (đã fix outline v3)',
         '3 (Q1-10, Q3-7, Q3-8)',
         'Em', 'ĐÃ FIX'),
        ('Tổng', '19', '—', '15 fixed + 4 chờ'),
    ],
    [4.0, 5.5, 3.5, 4.0]
)

H2('4.2 Outline v3 đã build với 15 fixes')
B('Outline_Q1_v3_01062026.docx — fix 8 issues + subtype-specific β table 7×3')
B('Outline_Q3_v3_01062026.docx — fix 7 issues + summary screening + grade '
  'trajectory figure')
B('OutlineBilingual_Q1_01062026.docx — bản song ngữ 15-18 trang gửi đồng '
  'tác giả')
B('OutlineBilingual_Q3_01062026.docx — bản song ngữ 13-15 trang gửi đồng '
  'tác giả')


# ============================================================
H1('5. 4 CÂU HỎI BLOCKING — CHỜ THẦY + NCS QUYẾT ĐỊNH')

H2('5.1 Bảng tổng hợp 4 quyết định cần thiết')
make_table(
    ['Mã', 'Vấn đề', 'Owner', 'Lựa chọn'],
    [
        ('Q1-6', 'Dữ liệu phỏng vấn định tính', 'NCS',
         'Số HS phỏng vấn? Sampling strategy? Transcripts? Cohen κ?'),
        ('Q1-8', 'R² mô hình SEM tích hợp', 'Thầy + NCS',
         '(A) Chạy lại integrated SEM với 7 dự báo → 1 R² tổng hợp; '
         'HOẶC (B) modify claim — "separate-model R² from parent dissertation"'),
        ('Q3-6', 'Văn bản phê duyệt đạo đức HNUE', 'NCS',
         'Đã có chính thức chưa? Nếu chưa, xin retroactive 2-4 tuần?'),
        ('Q3-9', 'Chiến lược tham chiếu chéo Q1 ↔ Q3', 'Thầy + NCS',
         '(A) Q1 first → Q3 cite Q1 (8-12 tháng); '
         'HOẶC (B) submit cùng lúc với "companion paper under review" '
         '(6 tháng)'),
    ],
    [1.2, 4.0, 2.5, 9.3]
)

H2('5.2 Hệ quả nếu không quyết sớm')
B('Q1-6 không có: claim "mixed-methods" không thể giữ → giảm 1/3 novel '
  'contributions')
B('Q1-8 không có integrated SEM: novel contribution 1 yếu → có thể bị '
  'reviewer reject')
B('Q3-6 không có IRB: BMC + PLOS ONE reject ngay từ editor screening')
B('Q3-9 không quyết: rủi ro self-plagiarism / redundant publication')


# ============================================================
H1('6. THAM KHẢO TITLES — CHÂU Á + CHÂU PHI')

H2('6.1 Tổng quan')
B('Đã tổng hợp 33+ titles từ 2023-2025 (18 Q1 + 15 Q3)')
B('Phạm vi: Trung Quốc, Việt Nam, Hàn Quốc, Bangladesh, Ấn Độ, Pakistan, '
  'Sri Lanka, Ethiopia, Nigeria, Rwanda, UAE, Qatar')
B('Nguồn: BMC Psychiatry, PLOS ONE, PLOS Mental Health, J Affective '
  'Disorders, Frontiers Psychiatry, Scientific Reports, General Psychiatry')

H2('6.2 Pattern phổ biến cho Q1')
B('SEM/Mediation + Population + Variable: "[Variable] and [Outcome] in '
  '[Population]: a structural equation modeling analysis"')
B('Multi-group invariance: "[Constructs] in [Population]: cross-gender '
  'measurement invariance of [scale]"')
B('Risk-Protective: "[Risk/Protective] factors of [Outcome] in [Population]"')

H2('6.3 Pattern phổ biến cho Q3')
B('Prevalence + AF (chuẩn nhất): "Prevalence and [associated factors / '
  'correlates / determinants] of [outcome] among [population] in [region]: '
  'a cross-sectional study"')
B('Subgroup comparison: "[Outcome] among [subgroup1, subgroup2, subgroup3] '
  'in [region]: Investigating prevalence and associated factors"')

H2('6.4 10 đề xuất tên cho nhóm chọn (5 Q1 + 5 Q3)')
P('Chi tiết đầy đủ trong file ThamKhao_Titles_Q1Q3_AsiaChauPhi_01062026.docx',
  italic=True)
B('Em đề xuất A4 cho Q1: "Multi-group structural equation modeling of risk '
  'and protective factors for anxiety disorder subtypes among Vietnamese '
  'adolescents: A mixed-methods cross-sectional study"')
B('Em đề xuất B2 cho Q3: "Prevalence and item-level patterns of generalized, '
  'separation, and social anxiety symptoms among Vietnamese lower secondary '
  'school students: A descriptive cross-sectional study"')


# ============================================================
H1('7. FILES QUAN TRỌNG ĐÃ TẠO')

H2('7.1 Outline + Bố cục Q1/Q3')
make_table(
    ['Tên file', 'Mục đích'],
    [
        ('BoCuc_Q1_Q3_PhuongAnA_v2_01062026.docx',
         'Bố cục tổng quan v2 (đã confirm chốt)'),
        ('Outline_Q1_v3_01062026.docx',
         'Outline Q1 v3 với 8 issues đã fix'),
        ('Outline_Q3_v3_01062026.docx',
         'Outline Q3 v3 với 7 issues đã fix'),
        ('OutlineBilingual_Q1_01062026.docx',
         'Bản song ngữ Q1 (gửi đồng tác giả)'),
        ('OutlineBilingual_Q3_01062026.docx',
         'Bản song ngữ Q3 (gửi đồng tác giả)'),
        ('ThamKhao_Titles_Q1Q3_AsiaChauPhi_01062026.docx',
         'Tham khảo 33+ titles + 10 đề xuất'),
    ],
    [8.0, 9.0]
)

H2('7.2 Báo cáo rà soát')
make_table(
    ['Tên file', 'Mục đích'],
    [
        ('RaSoat_Q1_01062026.docx',
         'Báo cáo rà soát Q1 với 10 issues'),
        ('RaSoat_Q3_01062026.docx',
         'Báo cáo rà soát Q3 với 9 issues'),
        ('IssuesPriority_Q1Q3_01062026.docx',
         'Phân loại 19 issues + timeline 2 tuần'),
        ('TomTat_Phien_01062026.docx',
         'Tài liệu này — tóm tắt toàn bộ công việc'),
    ],
    [8.0, 9.0]
)

H2('7.3 Tài liệu chính thức luận án TS')
make_table(
    ['Tên file', 'Mục đích'],
    [
        ('TomTatLA_v2_VERIFIED_29052026.docx',
         'Tóm tắt LA tiếng Việt'),
        ('TomTatLA_EN_v1_29052026.docx',
         'Tóm tắt LA tiếng Anh'),
        ('TrichYeuLA_CongThiHang_v2_29052026.docx',
         'Trích yếu LA song ngữ'),
        ('DanhMucCongTrinh_EN_v1_29052026.docx',
         'Danh mục công trình tiếng Anh'),
        ('Buoc5_GiayTiepNhan_PBDL_Lan1_DaDien_29052026.docx',
         'Form Bước 5 đã điền'),
    ],
    [8.0, 9.0]
)


# ============================================================
H1('8. BƯỚC TIẾP THEO ĐỀ XUẤT')

H2('8.1 Tuần 1 (01-07/06)')
B('Thầy + NCS quyết 4 BLOCKING questions')
B('NCS xin văn bản IRB từ HNUE (2-4 tuần — bắt đầu ngay)')
B('Thầy + đồng tác giả review 2 outline song ngữ Q1 + Q3')
B('Nhóm chọn tên bài từ 10 đề xuất (5 Q1 + 5 Q3) trong file Tham khảo')

H2('8.2 Tuần 2 (08-14/06)')
B('Em integrate quyết định từ thầy + NCS vào outline v3 FINAL')
B('Em viết đầy đủ ethics statement Q3 (sau khi có IRB letter)')
B('Em hoàn thiện cross-ref strategy Q1 ↔ Q3')
B('Sẵn sàng bắt đầu DRAFT Q1 tuần 3+')

H2('8.3 Các tuần tiếp')
B('Tuần 3-5: Em viết draft Q1 (Methods → Results → Discussion → Intro → Abstract)')
B('Tuần 6: Verify + plagiarism check + format BMC Psychiatry')
B('Tuần 7: Submit Q1 + cover letter + suggested reviewers')
B('Tuần 8+: Draft Q3 song song với chờ peer review Q1')


# Clean metadata
cp = d.core_properties
cp.author = ''; cp.title = ''; cp.subject = ''
cp.keywords = ''; cp.comments = ''; cp.last_modified_by = ''
cp.category = ''; cp.identifier = ''; cp.version = ''

d.save(OUT)
print(f'Saved: {OUT}')
print(f'Size: {os.path.getsize(OUT)} bytes')
