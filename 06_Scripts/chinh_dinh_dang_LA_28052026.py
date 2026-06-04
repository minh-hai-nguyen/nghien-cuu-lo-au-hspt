# -*- coding: utf-8 -*-
"""Chinh dinh dang LA theo chuan TS Viet Nam.
- Times New Roman 13pt, line 1.5, justify, indent 1.25cm
- Margins: Top 2.5 / Bottom 2.5 / Left 3.5 / Right 2.0 cm
- Headings: Chuong 16pt CAPS BOLD center, muc 14pt bold left, muc 3 13pt bold
- Insert TOC field at start
- Remove watermarks
- KHONG dong noi dung
28/05/2026."""
import os, sys, io, copy
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_6_FixAnderson_27052026.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v4_ChuanTrinhBay_28052026.docx')

print(f"SRC: {SRC}")
print(f"OUT: {OUT}")
print(f"SRC size: {os.path.getsize(SRC)//1024} KB")
print()


# ============================================================
# STEP 1: Open document
# ============================================================
d = Document(SRC)

# ============================================================
# STEP 2: Set section margins + remove headers/footers watermarks
# ============================================================
print("--- Step 2: Margins + watermark removal ---")
removed_watermarks = 0
for sec in d.sections:
    sec.top_margin = Cm(2.5)
    sec.bottom_margin = Cm(2.5)
    sec.left_margin = Cm(3.5)
    sec.right_margin = Cm(2.0)
    # Clear header/footer pict elements
    for hf in [sec.header, sec.first_page_header, sec.even_page_header,
               sec.footer, sec.first_page_footer, sec.even_page_footer]:
        try:
            for elem in list(hf._element.iter()):
                if elem.tag in (qn('w:pict'), qn('w:object'), qn('w:drawing')):
                    if elem.getparent() is not None:
                        elem.getparent().remove(elem)
                        removed_watermarks += 1
        except Exception as e:
            pass
print(f"  Margins set: 2.5/2.5/3.5/2.0 cm")
print(f"  Removed {removed_watermarks} watermark/picture elements from headers/footers")

# Also remove pict/drawing in body (floating watermark text boxes etc.)
body_watermarks = 0
body = d.element.body
for elem in list(body.iter()):
    tag = elem.tag
    # Remove w:pict (older watermarks) and check for textbox watermarks
    if tag == qn('w:pict'):
        # Check if this pict contains a watermark-like shape (textbox)
        # Watermarks usually have v:shape with type containing "PowerPlusWaterMarkObject"
        # For safety, remove pict only in headers (above) and skip body pict
        # But user wants ALL watermarks removed
        # Check shape type:
        is_watermark = False
        for sub in elem.iter():
            stype = sub.get('type', '') or sub.get(qn('o:spt'), '')
            if 'WaterMark' in stype or 'watermark' in str(stype).lower():
                is_watermark = True
                break
            # Check for text "Claude" or similar
            for txt_elem in sub.iter():
                if txt_elem.tag.endswith('}t') or txt_elem.tag.endswith('}textpath'):
                    text_content = txt_elem.text or ''
                    if any(kw in text_content.lower() for kw in ['claude', 'watermark', 'draft']):
                        is_watermark = True
                        break
            if is_watermark:
                break
        if is_watermark and elem.getparent() is not None:
            elem.getparent().remove(elem)
            body_watermarks += 1
print(f"  Removed {body_watermarks} watermark elements from body")
print()


# ============================================================
# STEP 3: Detect heading level for each paragraph
# ============================================================
def detect_heading_level(p):
    """Return heading level (1=chuong, 2=1.1, 3=1.1.1, 4=deeper) or 0 if body.
    Critical: only match short, ALL-CAPS, structural headings — NOT body text starting with 'Chương trình'."""
    import re
    text = p.text.strip()
    if not text:
        return 0
    # By style first (most reliable)
    style_name = p.style.name if p.style else ''
    if 'Heading 1' in style_name or 'Title' in style_name:
        return 1
    if 'Heading 2' in style_name:
        return 2
    if 'Heading 3' in style_name:
        return 3
    if 'Heading 4' in style_name or 'Heading 5' in style_name:
        return 4

    # Heading must be SHORT (heading text rarely > 200 chars)
    if len(text) > 200:
        return 0

    # Pattern: CHƯƠNG <number/roman> ... — must be ALL CAPS to avoid 'Chương trình', 'Chương 2 đã trình bày...'
    # Match: 'CHƯƠNG 1', 'CHƯƠNG 1:', 'CHƯƠNG 1. TITLE', 'CHƯƠNG I'
    if re.match(r'^CHƯƠNG\s+[\dIVX]+', text):
        # text starts with "CHƯƠNG <num>" in caps
        return 1

    # Other H1 keywords: must match WHOLE text or with colon/period after
    upper = text.upper()
    h1_keywords = ['MỞ ĐẦU', 'KẾT LUẬN', 'KIẾN NGHỊ', 'TÀI LIỆU THAM KHẢO',
                   'PHỤ LỤC', 'MỤC LỤC', 'LỜI CAM ĐOAN', 'LỜI CẢM ƠN',
                   'DANH MỤC CÔNG TRÌNH ĐÃ CÔNG BỐ',
                   'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
                   'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
                   'KẾT LUẬN VÀ KIẾN NGHỊ']
    for kw in h1_keywords:
        # Match if text equals kw, or kw followed by punctuation/end
        if upper == kw or upper.startswith(kw + ':') or upper.startswith(kw + '.'):
            return 1

    # Numbering pattern: only if SHORT (< 150 chars) — actual headings are short
    if len(text) < 150:
        if re.match(r'^\d+\.\d+\.\d+\.\d+\.?\s', text):
            return 4
        if re.match(r'^\d+\.\d+\.\d+\.?\s', text):
            return 3
        if re.match(r'^\d+\.\d+\.?\s', text):
            return 2
    return 0


# ============================================================
# STEP 4: Apply paragraph + run formatting
# ============================================================
print("--- Step 4: Apply paragraph formatting ---")
n_total = 0
n_h1 = n_h2 = n_h3 = n_h4 = n_body = 0
for p in d.paragraphs:
    n_total += 1
    level = detect_heading_level(p)
    pf = p.paragraph_format
    # Reset
    pf.line_spacing = 1.5
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.space_before = Pt(0)
    if level == 1:
        # Chương: 16pt UPPER BOLD center, no indent
        pf.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(18)
        pf.space_after = Pt(12)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(16)
            r.bold = True
        n_h1 += 1
    elif level == 2:
        # 1.1.: 14pt bold left, no indent
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(12)
        pf.space_after = Pt(6)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(14)
            r.bold = True
        n_h2 += 1
    elif level == 3:
        # 1.1.1.: 13pt bold left, no indent
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(10)
        pf.space_after = Pt(4)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            r.bold = True
        n_h3 += 1
    elif level == 4:
        # 1.1.1.1.: 13pt italic bold
        pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf.first_line_indent = Cm(0)
        pf.space_before = Pt(6)
        pf.space_after = Pt(3)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            r.bold = True
            r.italic = True
        n_h4 += 1
    else:
        # Body: 13pt justify, indent 1.25cm
        pf.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        # Preserve indent ONLY for body that already had it (avoid breaking list items)
        if pf.first_line_indent is None or pf.first_line_indent.cm < 0.5:
            pf.first_line_indent = Cm(1.25)
        pf.space_after = Pt(6)
        for r in p.runs:
            r.font.name = 'Times New Roman'
            r.font.size = Pt(13)
            # Preserve bold/italic/color (don't touch)
        n_body += 1

print(f"  Total paragraphs: {n_total}")
print(f"  H1 (chuong): {n_h1} | H2: {n_h2} | H3: {n_h3} | H4: {n_h4} | Body: {n_body}")
print()


# ============================================================
# STEP 5: Apply formatting to tables
# ============================================================
print("--- Step 5: Apply table formatting ---")
n_table_cells = 0
for tb in d.tables:
    for row in tb.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                pf = p.paragraph_format
                pf.line_spacing = 1.15
                pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
                pf.space_before = Pt(0)
                pf.space_after = Pt(0)
                # Tables: no first line indent
                pf.first_line_indent = Cm(0)
                for r in p.runs:
                    r.font.name = 'Times New Roman'
                    # Keep cell text smaller than body (12pt)
                    if r.font.size is None or r.font.size.pt > 13:
                        r.font.size = Pt(12)
                n_table_cells += 1
print(f"  Total table cells processed: {n_table_cells}")
print()


# ============================================================
# STEP 6: Insert TOC field at the beginning
# ============================================================
print("--- Step 6: Insert TOC field ---")

def make_toc_paragraph(parent_para):
    """Insert TOC field code into a paragraph."""
    r = parent_para.add_run()
    # Begin field
    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    r._r.append(fldChar_begin)
    # Field instruction
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = r'TOC \o "1-3" \h \z \u'
    r._r.append(instrText)
    # Separator
    fldChar_sep = OxmlElement('w:fldChar')
    fldChar_sep.set(qn('w:fldCharType'), 'separate')
    r._r.append(fldChar_sep)
    # Result text placeholder
    instrResult = OxmlElement('w:t')
    instrResult.text = 'Nhấn F9 hoặc chuột phải > Update Field để cập nhật Mục lục.'
    r._r.append(instrResult)
    # End field
    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    r._r.append(fldChar_end)


# Check if Mục lục already exists (scan first 200 paragraphs - cover + front matter)
toc_exists = False
toc_para_idx = None
for i, p in enumerate(d.paragraphs[:200]):
    text = p.text.strip().upper()
    if text in ('MỤC LỤC', 'MUC LUC'):
        toc_exists = True
        toc_para_idx = i
        break

if toc_exists:
    print(f"  Tim thay tieu de 'MỤC LỤC' tai para {toc_para_idx}. Chen TOC ngay sau.")
    # Insert new paragraph for TOC after the heading
    toc_heading_p = d.paragraphs[toc_para_idx]
    # Format the heading as Heading 1
    toc_heading_p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    toc_heading_p.paragraph_format.first_line_indent = Cm(0)
    for r in toc_heading_p.runs:
        r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    # Insert TOC after this paragraph
    new_p = OxmlElement('w:p')
    toc_heading_p._element.addnext(new_p)
    # Now add field to new_p via wrapping
    from docx.text.paragraph import Paragraph
    toc_p_obj = Paragraph(new_p, toc_heading_p._parent)
    make_toc_paragraph(toc_p_obj)
else:
    print("  Khong tim thay MỤC LỤC. Chen tieu de + TOC vao dau document.")
    # Get the first paragraph element
    first_p = d.paragraphs[0]._element
    body = first_p.getparent()
    # Create TOC heading
    heading_p = OxmlElement('w:p')
    body.insert(list(body).index(first_p), heading_p)
    from docx.text.paragraph import Paragraph
    heading_obj = Paragraph(heading_p, d.paragraphs[0]._parent)
    r = heading_obj.add_run('MỤC LỤC')
    r.font.name = 'Times New Roman'; r.font.size = Pt(16); r.bold = True
    heading_obj.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    heading_obj.paragraph_format.space_before = Pt(18)
    heading_obj.paragraph_format.space_after = Pt(12)
    # Insert TOC paragraph after heading
    toc_p_xml = OxmlElement('w:p')
    body.insert(list(body).index(heading_p) + 1, toc_p_xml)
    toc_p_obj = Paragraph(toc_p_xml, d.paragraphs[0]._parent)
    make_toc_paragraph(toc_p_obj)
    # Insert page break after TOC
    pb_p = OxmlElement('w:p')
    body.insert(list(body).index(toc_p_xml) + 1, pb_p)
    pb_obj = Paragraph(pb_p, d.paragraphs[0]._parent)
    pb_r = pb_obj.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    pb_r._r.append(br)

print()


# ============================================================
# STEP 7: Clean core properties (metadata)
# ============================================================
print("--- Step 7: Clean metadata ---")
cp = d.core_properties
cp.author = ''
cp.title = ''
cp.subject = ''
cp.keywords = ''
cp.comments = ''
cp.last_modified_by = ''
cp.category = ''
print("  Metadata cleared")
print()


# ============================================================
# STEP 8: Save
# ============================================================
d.save(OUT)
print(f"--- Done ---")
print(f"Saved: {OUT}")
print(f"Size: {os.path.getsize(OUT)//1024} KB")
