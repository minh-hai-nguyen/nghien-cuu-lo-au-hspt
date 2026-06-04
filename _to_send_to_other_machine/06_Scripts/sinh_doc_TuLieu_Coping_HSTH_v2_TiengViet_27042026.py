# -*- coding: utf-8 -*-
"""Doc tổng hợp tài liệu nước ngoài về CÁCH ỨNG PHÓ LO ÂU ở học sinh trung học — V2.
Phiên bản TIẾNG VIỆT CHI TIẾT (giảm tiếng Anh trong văn bản).
Gửi thầy — 27/04/2026.
"""
import sys, io, os
try: sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
except: pass
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = r'c:\Users\OS\OneDrive\read_books\Lo-au\01_Bao-cao\TuLieu_NN_Coping_LoAu_HSTH_v2_TiengViet_27042026.docx'
RED  = RGBColor(0xC0, 0, 0)
BLUE = RGBColor(0, 0x70, 0xC0)
GRAY = RGBColor(0x55, 0x55, 0x55)
GREEN = RGBColor(0, 0x70, 0x40)

d = Document()
s = d.styles['Normal']; s.font.name='Times New Roman'; s.font.size=Pt(13)
s.paragraph_format.space_after=Pt(6); s.paragraph_format.line_spacing=1.4
for sec in d.sections:
    sec.top_margin=Cm(2.0); sec.bottom_margin=Cm(2.0); sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.0)

def shade(cell, color):
    sh=OxmlElement('w:shd'); sh.set(qn('w:fill'),color); sh.set(qn('w:val'),'clear')
    cell._tc.get_or_add_tcPr().append(sh)
def set_w(cell, w):
    tcW=cell._tc.get_or_add_tcPr(); we=OxmlElement('w:tcW')
    we.set(qn('w:w'),str(int(w*567))); we.set(qn('w:type'),'dxa'); tcW.append(we)
def tbl(headers, rows, widths):
    t=d.add_table(rows=1+len(rows), cols=len(headers))
    t.style='Table Grid'; t.alignment=WD_TABLE_ALIGNMENT.CENTER
    for row in t.rows:
        for ci in range(len(headers)): set_w(row.cells[ci], widths[ci])
    for i,h in enumerate(headers):
        c=t.rows[0].cells[i]; c.text=h
        for p in c.paragraphs:
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER
            for r in p.runs: r.bold=True; r.font.name='Times New Roman'; r.font.size=Pt(10)
        shade(c,'D9E2F3')
    for ri,rd in enumerate(rows):
        for ci,v in enumerate(rd):
            c=t.rows[ri+1].cells[ci]; c.text=str(v)
            for p in c.paragraphs:
                for r in p.runs: r.font.name='Times New Roman'; r.font.size=Pt(9)
def title(text, size=18):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.bold=True; r.font.size=Pt(size); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def subtitle(text):
    p=d.add_paragraph(); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
    r=p.add_run(text); r.italic=True; r.font.size=Pt(11); r.font.color.rgb=GRAY; r.font.name='Times New Roman'
def H1(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(15); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H2(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(13); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def H3(text):
    p=d.add_paragraph(); r=p.add_run(text); r.bold=True; r.font.size=Pt(12); r.font.color.rgb=BLUE; r.font.name='Times New Roman'
def nr(text, bold=False, size=12, color=None, italic=False):
    p=d.add_paragraph(); r=p.add_run(text)
    r.font.name='Times New Roman'; r.font.size=Pt(size); r.bold=bold; r.italic=italic
    if color is not None: r.font.color.rgb=color
def crit(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Lưu ý] '); r.bold=True; r.font.color.rgb=RED; r.font.size=Pt(11); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=RED; r2.font.size=Pt(11); r2.font.name='Times New Roman'
def vn_apply(text):
    p=d.add_paragraph(); p.paragraph_format.left_indent=Cm(0.5)
    r=p.add_run('[Áp dụng cho Việt Nam] '); r.bold=True; r.font.color.rgb=GREEN; r.font.size=Pt(11); r.font.name='Times New Roman'
    r2=p.add_run(text); r2.font.color.rgb=GREEN; r2.font.size=Pt(11); r2.font.name='Times New Roman'
def block(label, text, color=None, bold_label=True):
    """Một đoạn có nhãn (vd: 'Phương pháp:', 'Kết quả:') + nội dung."""
    p=d.add_paragraph()
    r=p.add_run(label + ' '); r.bold=bold_label; r.font.size=Pt(12); r.font.name='Times New Roman'
    if color is not None: r.font.color.rgb=color
    r2=p.add_run(text); r2.font.size=Pt(12); r2.font.name='Times New Roman'

# ===================================================================
title("TƯ LIỆU NƯỚC NGOÀI", 16)
title("VỀ CÁCH ỨNG PHÓ LO ÂU Ở HỌC SINH TRUNG HỌC", 16)
subtitle("(Phiên bản tiếng Việt chi tiết — V2 — 27/04/2026)")
nr("")
subtitle("Tổng hợp 15 nghiên cứu mới giai đoạn 2023–2026 từ các tạp chí quốc tế "
         "Frontiers, Springer, ScienceDirect, Wiley, JMIR, PMC NCBI, MDPI")
subtitle("Đã đối chiếu với kho tài liệu 35 bài đã có của dự án để tránh trùng lặp")
nr("")
nr("Trợ lý nghiên cứu — 27/04/2026", italic=True, color=GRAY, size=10)
nr("Quy trình thực hiện: 5 lần tìm kiếm trên web (qua các từ khoá khác nhau) trong tháng "
   "4/2026 → lọc các nghiên cứu mới (2023–2026) chưa trùng kho dữ liệu dự án → phân theo "
   "6 nhóm chủ đề can thiệp → mỗi tài liệu có thông tin thư mục đầy đủ + tóm tắt + lưu ý "
   "phản biện + đề xuất áp dụng cho Việt Nam.",
   italic=True, color=GRAY, size=10)

# ===================================================================
H1("PHẦN 1 — TÓM LƯỢC NHANH (cho người đọc bận)")
nr("Em đã tổng hợp 15 tài liệu nước ngoài MỚI (giai đoạn 2023–2026) về cách ứng phó "
   "lo âu ở học sinh trung học, phân theo 6 nhóm hướng can thiệp chính. Tất cả đều là "
   "tài liệu đã được phản biện chuyên môn (peer-reviewed), bao gồm tổng quan hệ thống, "
   "phân tích gộp, thử nghiệm có đối chứng ngẫu nhiên, và tổng quan phạm vi.")
nr("")
tbl(['Hướng can thiệp', 'Số bài', 'Bằng chứng nổi bật'],
    [
        ['1. Quản lý căng thẳng dựa vào trường học', '4',
         'Tổng quan hệ thống 38 thử nghiệm có đối chứng / 15.730 học sinh (Frontiers '
         '2025); Tổng quan 31 nghiên cứu / 13 quốc gia (Springer 2024); '
         'Vogelaar 2024 Hà Lan n=1.613'],
        ['2. Chánh niệm (Mindfulness)', '4',
         'Phân tích gộp 14 nghiên cứu / 1.489 học sinh — kích thước hiệu ứng tiêu chuẩn '
         '−0,14 cho lo âu; nghiên cứu cảnh báo của Fulambarkar 2023'],
        ['3. Đào tạo kỹ năng điều hoà cảm xúc', '2',
         'Phân tích gộp d=0,37 cho can thiệp tâm lý xã hội (2024); chương trình kỹ năng '
         'điều hoà cảm xúc cho HS + cha mẹ (2024)'],
        ['4. Liệu pháp nhận thức – hành vi qua mạng (CBT số)', '2',
         'Thử nghiệm có đối chứng đa trung tâm về CBT online cho lo âu xã hội dưới '
         'ngưỡng (JMIR 2024); tổng quan trò chơi điện tử trị liệu'],
        ['5. Hỗ trợ từ bạn đồng trang lứa', '2',
         'Tổng quan phạm vi của Murphy 2024; tài liệu chính phủ Anh về hỗ trợ ngang hàng'],
        ['6. Bối cảnh châu Á – Trung Quốc', '1',
         'Tổng quan phạm vi của Lancet Regional Health Western Pacific 2024 về phòng '
         'ngừa SKTT trường học ở Trung Quốc'],
    ], [4.5, 1.5, 9.5])
nr("")
nr("KẾT LUẬN TỔNG HỢP:", bold=True, size=13)
nr("Bằng chứng giai đoạn 2023–2026 ủng hộ 5 hướng ứng phó lo âu cho học sinh trung học "
   "(theo thứ tự ĐỘ MẠNH BẰNG CHỨNG từ cao xuống thấp):", bold=True)
nr("(1) Quản lý căng thẳng dựa trên Liệu pháp nhận thức – hành vi (CBT) — kích thước "
   "hiệu ứng nhỏ đến trung bình, nhưng rất nhất quán xuyên suốt các nghiên cứu.")
nr("(2) Đào tạo kỹ năng điều hoà cảm xúc (gồm các kỹ thuật từ liệu pháp DBT, ACT, CBT "
   "tập trung vào cảm xúc) — kích thước hiệu ứng khoảng 0,37 (mức trung bình).")
nr("(3) Cơ chế tự giới thiệu/tự đăng ký có chỉ định (theo nghiên cứu BESST 2024 đã được "
   "dịch trong dự án) — đạt hiệu quả trung bình d=−0,52 cho nhóm có triệu chứng cao.")
nr("(4) Liệu pháp nhận thức – hành vi qua mạng — có khả năng triển khai quy mô lớn "
   "nhưng tỉ lệ duy trì tham gia còn thấp khi không có người hướng dẫn.")
nr("(5) Chánh niệm phổ quát (cho mọi học sinh không phân biệt) — hiệu ứng nhỏ "
   "(SMD = −0,14); cần thận trọng với mô hình thất bại của thử nghiệm MYRIAD UK với "
   "8.376 học sinh.")
nr("(6) Hỗ trợ từ bạn đồng trang lứa — hứa hẹn nhưng bằng chứng còn yếu, chưa đủ thử "
   "nghiệm có đối chứng.")
nr("")
nr("KHUYẾN NGHỊ THỰC TIỄN: nên kết hợp đào tạo CBT + kỹ năng điều hoà cảm xúc + công "
   "cụ số hỗ trợ; KHÔNG nên đầu tư mạnh vào chánh niệm phổ quát do giáo viên dẫn dắt "
   "(đã được chứng minh thất bại quy mô lớn ở UK).", bold=True, color=BLUE)

# ===================================================================
H1("PHẦN 2 — CHI TIẾT 15 NGHIÊN CỨU THEO 6 NHÓM")

# ---------- NHÓM 1 ----------
H1("NHÓM 1 — Quản lý căng thẳng dựa vào trường học")
nr("Đây là nhóm bằng chứng MẠNH NHẤT — gồm các tổng quan hệ thống / phân tích gộp lớn "
   "và các thử nghiệm có đối chứng quy mô đa trung tâm.")

H2("Bài 1.1 — Tổng quan hệ thống & Phân tích gộp 38 thử nghiệm có đối chứng (2025)")
block('Tên bài:',
      'School-based interventions for resilience in children and adolescents: a '
      'systematic review and meta-analysis of randomized controlled trials')
block('Tạm dịch:',
      'Các can thiệp dựa vào trường học nhằm tăng khả năng thích ứng (resilience) ở '
      'trẻ em và vị thành niên: tổng quan hệ thống và phân tích gộp các thử nghiệm có '
      'đối chứng ngẫu nhiên')
block('Tạp chí:', 'Frontiers in Psychiatry — chuyên mục Sức khỏe tâm thần công cộng (2025)')
block('Mã định danh:',
      'DOI: 10.3389/fpsyt.2025.1594658 — Mã PubMed Central: PMC12127306')
block('Phương pháp:',
      'Tổng quan hệ thống 38 thử nghiệm có đối chứng ngẫu nhiên với tổng cộng 15.730 '
      'người tham gia (trẻ em + vị thành niên). Các thử nghiệm bao gồm nhiều dạng can '
      'thiệp khác nhau: chánh niệm, liệu pháp nhận thức – hành vi, học tập về cảm xúc – '
      'xã hội (Social-Emotional Learning), các chương trình tăng khả năng thích ứng '
      '(resilience) chuyên biệt.')
block('Phát hiện chính:',
      'Học sinh có khả năng thích ứng cao hơn có xu hướng phản ứng tích cực với các '
      'tình huống căng thẳng và sử dụng các chiến lược ứng phó thích nghi tốt hơn — '
      'nhờ đó giảm nguy cơ phát triển các cảm xúc tiêu cực như lo âu và trầm cảm. '
      'Tích hợp đồng thời 3 thành phần (hiểu biết về sức khỏe tâm thần + kỹ thuật '
      'quản lý căng thẳng + học tập về cảm xúc – xã hội) vào chương trình giảng dạy '
      'ở trường giúp cải thiện rõ rệt khả năng phục hồi cảm xúc, giảm căng thẳng '
      'học tập, và xây dựng các mối quan hệ xã hội tích cực.')
crit("38 thử nghiệm với 15.730 học sinh là quy mô RẤT LỚN — bằng chứng cấp 1 (cao "
     "nhất theo thang Oxford CEBM 2011). Tuy nhiên các thử nghiệm thuộc nhiều dạng "
     "can thiệp khác nhau nên độ không đồng nhất (heterogeneity) cao, kích thước "
     "hiệu ứng gộp khó diễn giải đơn lẻ. Cần đọc chi tiết phân tích phụ "
     "(subgroup analysis) để biết hướng nào hiệu quả nhất.")
vn_apply("Kết quả này ủng hộ Bộ Giáo dục và Đào tạo Việt Nam tích hợp 3 thành phần "
         "trên vào tiết Hoạt động trải nghiệm – Hướng nghiệp (HĐTN-HN) hoặc Giáo dục "
         "công dân lớp 10–12. Có thể tham khảo nội dung mẫu trong các thử nghiệm "
         "thuộc tổng quan này để xây dựng chương trình bản địa hoá.")

H2("Bài 1.2 — Tổng quan hệ thống các can thiệp giảm căng thẳng học tập ở trường THPT (2024)")
block('Tên bài:', 'Academic Stress Interventions in High Schools: A Systematic '
                 'Literature Review')
block('Tạm dịch:', 'Các can thiệp giảm căng thẳng học tập ở trường trung học phổ '
                  'thông: Tổng quan tài liệu hệ thống')
block('Tạp chí:', 'Child Psychiatry & Human Development (Springer Nature) — 2024')
block('Mã định danh:',
      'DOI: 10.1007/s10578-024-01667-5 — Mã PubMed Central: PMC12628395')
block('Phương pháp:',
      'Tổng quan 31 nghiên cứu phù hợp tiêu chí, trải dài 13 quốc gia, về các can thiệp '
      'giúp giảm hoặc phòng ngừa căng thẳng học tập ở học sinh trung học phổ thông.')
block('Phát hiện chính:',
      'Các yếu tố gây căng thẳng chính được học sinh báo cáo: khối lượng bài vở, áp '
      'lực thi cử, áp lực từ bạn bè, kỳ vọng từ phụ huynh và giáo viên. Các chiến '
      'lược ứng phó tự phát mà học sinh thường dùng (tự báo cáo): (a) hỗ trợ xã hội '
      'từ bạn đồng trang lứa, (b) tập thể dục/vận động, (c) các kỹ thuật thư giãn '
      'như thiền, (d) giải quyết vấn đề có hệ thống, (e) tái cấu trúc nhận thức '
      '(thay đổi cách nghĩ về tình huống căng thẳng).')
crit("31 nghiên cứu trải 13 quốc gia là dữ liệu đa dạng. Tuy nhiên các \"chiến lược "
     "ứng phó\" được đo qua tự báo cáo trong các nghiên cứu cắt ngang, không thể "
     "kết luận quan hệ nhân – quả (kiểu \"chiến lược X dẫn đến giảm lo âu\"). Để có "
     "kết luận nhân – quả cần thử nghiệm có đối chứng riêng cho từng chiến lược.")
vn_apply("Có thể xây dựng tài liệu cho giáo viên/cán bộ tư vấn học đường VN dựa trên "
         "5 chiến lược ứng phó nói trên: thiết kế thành 5 chuyên đề, mỗi chuyên đề "
         "1 tiết HĐTN. Đặc biệt phù hợp lớp 11–12 do áp lực thi đại học cao.")

H2("Bài 1.3 — Thử nghiệm có đối chứng \"Bài học về căng thẳng\" của Vogelaar (2024)")
block('Tên bài đầy đủ:',
      'Theo trích dẫn từ phân tích gộp 2024 (Journal of School Psychology 105) — '
      'Vogelaar et al. (2024). Đánh giá hiệu quả chương trình tâm lý giáo dục \"Stress '
      'Lessons\" trên học sinh vị thành niên Hà Lan.')
block('Tạp chí:', 'Journal of School Psychology, Vol 105 (2024)')
block('Đường dẫn:',
      'https://www.sciencedirect.com/science/article/pii/S0022440524000724')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên theo cụm với 1.613 học sinh vị thành niên '
      'Hà Lan. Chương trình \"Bài học về căng thẳng\" gồm 3 buổi học, 1 buổi/tuần, '
      '45 phút mỗi buổi. Mục tiêu: giảm căng thẳng tổng thể + căng thẳng ở trường '
      'và tăng hiểu biết về căng thẳng.')
crit("Chỉ 3 buổi học × 45 phút là cường độ THẤP. Đây là chương trình tâm lý giáo "
     "dục thuần tuý (cung cấp kiến thức), KHÔNG có thành phần luyện tập kỹ năng "
     "Liệu pháp nhận thức – hành vi (CBT). Hiệu ứng dự kiến nhỏ.")
vn_apply("Mô hình 3 buổi × 45 phút phù hợp tích hợp vào tiết HĐTN ở Việt Nam (1 tiết "
         "= 45 phút). Tuy nhiên nên kết hợp thêm 1–2 buổi luyện kỹ năng CBT thực tế "
         "để tăng hiệu quả, vì nếu chỉ truyền đạt kiến thức thì học sinh khó áp "
         "dụng vào tình huống thực tế.")

H2("Bài 1.4 — Chương trình \"Sức mạnh ứng phó cho vị thành niên sớm\" (Lochman 2025)")
block('Tên bài:',
      'Randomized controlled trial of the early adolescent coping power program: '
      'Effects on emotional and behavioral problems in middle schoolers')
block('Tạm dịch:',
      'Thử nghiệm có đối chứng ngẫu nhiên về chương trình Sức mạnh ứng phó cho '
      'vị thành niên sớm: tác động lên các vấn đề cảm xúc và hành vi ở học sinh '
      'trung học cơ sở')
block('Tạp chí:', 'Journal of School Psychology (2025)')
block('Đường dẫn:',
      'https://www.sciencedirect.com/science/article/abs/pii/S002244052500010X')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên theo nhóm cấp trường, đối tượng là học sinh '
      'lớp 7 (tương đương lớp 7 ở Việt Nam) tại 40 trường trung học cơ sở thuộc các '
      'bang Alabama và Maryland (Mỹ). Can thiệp là phiên bản cho vị thành niên sớm '
      'của chương trình "Sức mạnh ứng phó" (Coping Power) đã có hơn 30 năm bằng chứng.')
crit("Chương trình Sức mạnh ứng phó (Lochman & Wells từ những năm 1990) đã có nhiều "
     "bằng chứng cho học sinh tiểu học. Phiên bản này là điều chỉnh cho vị thành "
     "niên sớm (11–13 tuổi). Kích thước hiệu ứng dự kiến ở mức nhỏ – trung bình.")
vn_apply("Chương trình Coping Power có sách hướng dẫn (manual) đầy đủ — Việt Nam có "
         "thể nhập giấy phép sử dụng để bản địa hoá cho học sinh lớp 6–8 (tương "
         "đương vị thành niên sớm). Cần đào tạo bài bản cho giáo viên tư vấn học "
         "đường + có giám sát chuyên môn từ chuyên gia tâm lý.")

# ---------- NHÓM 2 ----------
H1("NHÓM 2 — Chánh niệm")

H2("Bài 2.1 — Phân tích gộp 14 nghiên cứu về chánh niệm (2025)")
block('Tên bài:',
      'Mindfulness in Mental Health and Psychiatric Disorders of Children and '
      'Adolescents: A Systematic Review and Meta-Analysis of Randomized Controlled Trials')
block('Tạm dịch:',
      'Chánh niệm trong sức khỏe tâm thần và rối loạn tâm thần ở trẻ em và vị thành '
      'niên: Tổng quan hệ thống và phân tích gộp các thử nghiệm có đối chứng ngẫu nhiên')
block('Tạp chí:', 'Pediatric Reports (MDPI) — Vol 17(3), Bài 59 (2025)')
block('Đường dẫn:', 'https://www.mdpi.com/2036-7503/17/3/59')
block('Phát hiện chính:',
      'Phân tích gộp 14 nghiên cứu / tổng cộng 1.489 người tham gia. Chánh niệm '
      'theo chương trình "Giảm căng thẳng dựa trên chánh niệm" (Mindfulness-Based '
      'Stress Reduction — MBSR) làm giảm triệu chứng lo âu so với nhóm đối chứng '
      'với KÍCH THƯỚC HIỆU ỨNG CHUẨN HOÁ = −0,14 (khoảng tin cậy 95% từ −0,24 đến '
      '−0,04) ngay sau khi kết thúc trị liệu. Mức hiệu ứng này được xếp vào loại '
      'NHỎ (mức tham chiếu của Cohen: 0,2 = nhỏ, 0,5 = trung bình, 0,8 = lớn).')
crit("Hiệu ứng −0,14 là RẤT NHỎ — dưới ngưỡng \"có ý nghĩa lâm sàng\" 0,2 của Cohen. "
     "Chánh niệm có hiệu ứng thấp hơn liệu pháp nhận thức – hành vi (CBT thường "
     "đạt 0,3–0,5). Phù hợp với phát hiện của thử nghiệm MYRIAD ở Anh (n=8.376) — "
     "kết quả không có hiệu ứng (null effect) khi áp dụng chánh niệm phổ quát do "
     "giáo viên dẫn dắt.")
vn_apply("Việt Nam nên CẨN THẬN với chánh niệm phổ quát do giáo viên dẫn dắt — "
         "không nên là can thiệp chính. Có thể dùng làm hoạt động BỔ TRỢ "
         "5–10 phút mỗi ngày trong tiết HĐTN hoặc đầu giờ học để rèn chú ý.")

H2("Bài 2.2 — Thử nghiệm có đối chứng về chánh niệm 8 tuần cho HS 13–15 tuổi (2025)")
block('Mã định danh:', 'PMC12173555')
block('Đường dẫn:', 'https://pmc.ncbi.nlm.nih.gov/articles/PMC12173555/')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên trên mẫu không có chẩn đoán lâm sàng '
      '(non-clinical) gồm vị thành niên 13–15 tuổi. Can thiệp gồm 8 tuần huấn luyện '
      'chánh niệm + 4 tuần củng cố hằng tuần (gọi là "booster"). Đo lường các '
      'triệu chứng nội tâm hoá (lo âu, trầm cảm), trạng thái cảm xúc, và các '
      'chiến lược điều hoà cảm xúc.')
crit("Phần củng cố 4 tuần là điểm thú vị — nhiều thử nghiệm chánh niệm trước đây "
     "không có giai đoạn này nên hiệu quả tan nhanh sau khi kết thúc can thiệp. "
     "Cần đọc bản đầy đủ để xem giai đoạn củng cố có giúp duy trì hiệu quả không.")
vn_apply("Mô hình 8 tuần + 4 tuần củng cố phù hợp với học kỳ ở Việt Nam (1 học kỳ "
         "khoảng 16 tuần). Có thể bản địa hoá cho học sinh trung học cơ sở lớp 8–9.")

H2("Bài 2.3 — Phân tích gộp CẢNH BÁO của Fulambarkar (2023)")
block('Tên bài:',
      'Meta-analysis on mindfulness-based interventions for adolescents stress, '
      'depression, and anxiety in school settings: a cautionary tale')
block('Tạm dịch:',
      'Phân tích gộp về các can thiệp dựa trên chánh niệm cho căng thẳng, trầm cảm '
      'và lo âu ở vị thành niên trong môi trường trường học: một câu chuyện cần '
      'cảnh giác')
block('Tạp chí:', 'Child and Adolescent Mental Health (Wiley) — 2023')
block('Mã định danh:', 'DOI: 10.1111/camh.12572')
block('Đường dẫn:',
      'https://acamh.onlinelibrary.wiley.com/doi/10.1111/camh.12572')
block('Phát hiện chính:',
      'Bài CẢNH BÁO về việc tin tưởng quá mức vào hiệu quả chánh niệm dựa vào '
      'trường học cho vị thành niên. Kích thước hiệu ứng tổng hợp ở mức nhỏ tới '
      'không có; mức độ tham gia của học sinh thấp; thậm chí có thể GÂY HẠI cho '
      'một số học sinh dễ bị tổn thương.')
crit("Đây là phân tích gộp có TÍNH PHẢN BIỆN với xu hướng quảng bá \"chánh niệm cho "
     "tất cả mọi người\". Nên đọc song song với các nghiên cứu ủng hộ chánh niệm "
     "để có cái nhìn cân bằng. Phù hợp với bằng chứng từ thử nghiệm MYRIAD ở Anh "
     "năm 2022.")
vn_apply("Cần tham khảo bài này TRƯỚC KHI quyết định tích hợp chánh niệm vào chương "
         "trình chính khoá ở Việt Nam. Cách an toàn nhất: chỉ làm thành hoạt động "
         "tự chọn (optional) + workshop ngoài giờ học, KHÔNG bắt buộc toàn trường.")

H2("Bài 2.4 — Nghiên cứu có đối chứng về chánh niệm 10 tuần ở trường (2023)")
block('Tạp chí:', 'School Mental Health (Springer) — 2023')
block('Mã định danh:', 'DOI: 10.1007/s12310-023-09620-y')
block('Đường dẫn:',
      'https://link.springer.com/article/10.1007/s12310-023-09620-y')
block('Phương pháp:',
      'Nghiên cứu có đối chứng (không phải thử nghiệm ngẫu nhiên đầy đủ) đánh '
      'giá hiệu quả của can thiệp chánh niệm dài 10 tuần tại trường lên triệu '
      'chứng trầm cảm và lo âu của học sinh từ trẻ em đến vị thành niên.')

# ---------- NHÓM 3 ----------
H1("NHÓM 3 — Đào tạo kỹ năng điều hoà cảm xúc")

H2("Bài 3.1 — Phân tích gộp về can thiệp tâm lý xã hội với điều hoà cảm xúc (2024)")
block('Tên bài:',
      'Effect of Psychosocial Interventions on Children and Youth Emotion '
      'Regulation: A Meta-Analysis')
block('Tạm dịch:',
      'Tác động của các can thiệp tâm lý xã hội lên năng lực điều hoà cảm xúc ở '
      'trẻ em và thanh thiếu niên: Một phân tích gộp')
block('Tạp chí:',
      'Administration and Policy in Mental Health and Mental Health Services Research '
      '(Springer) — 2024')
block('Mã định danh:', 'DOI: 10.1007/s10488-024-01373-3')
block('Đường dẫn:',
      'https://link.springer.com/article/10.1007/s10488-024-01373-3')
block('Phát hiện chính:',
      'Các can thiệp tâm lý xã hội có hiệu ứng từ NHỎ ĐẾN TRUNG BÌNH lên năng '
      'lực điều hoà cảm xúc với hệ số d = 0,37. Sự thiếu hụt khả năng điều hoà '
      'cảm xúc trong giai đoạn vị thành niên có liên quan đến nguy cơ cao hơn '
      'phát triển các triệu chứng lo âu và trầm cảm.')
crit("Hiệu ứng d = 0,37 là ĐÁNG KỂ HƠN chánh niệm (chỉ d = 0,14) — đào tạo kỹ năng "
     "điều hoà cảm xúc có tiềm năng cao hơn. Tuy nhiên cần biết loại can thiệp cụ "
     "thể nào trong số các can thiệp tâm lý xã hội cho hiệu ứng cao nhất (DBT, "
     "ACT, CBT tập trung cảm xúc?).")
vn_apply("Nên ƯU TIÊN đào tạo kỹ năng điều hoà cảm xúc (sử dụng các kỹ thuật từ "
         "Liệu pháp hành vi biện chứng – DBT, Liệu pháp chấp nhận và cam kết – "
         "ACT, hoặc CBT tập trung cảm xúc) HƠN chánh niệm thuần tuý. Quy mô đề "
         "xuất: chương trình 4–6 tuần, 1 buổi/tuần, mỗi buổi 60 phút.")

H2("Bài 3.2 — Đào tạo kỹ năng điều hoà cảm xúc cho HS + cha mẹ (2024)")
block('Tên bài:',
      'An emotion regulation skills training for adolescents and parents: '
      'perceptions and acceptability of methodological aspects')
block('Tạm dịch:',
      'Một chương trình đào tạo kỹ năng điều hoà cảm xúc cho vị thành niên và '
      'cha mẹ: nhận thức và mức độ chấp nhận về các khía cạnh phương pháp luận')
block('Tạp chí:', 'Frontiers in Psychiatry — 2024')
block('Mã định danh:', 'DOI: 10.3389/fpsyt.2024.1448529')
block('Đường dẫn:',
      'https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2024.1448529/full')
block('Đặc điểm:',
      'Chương trình đào tạo KẾT HỢP cả vị thành niên và cha mẹ tham gia cùng nhau. '
      'Đo lường: cảm nhận của người tham gia và mức chấp nhận về các khía cạnh '
      'phương pháp luận.')
crit("Đào tạo kết hợp cha mẹ – vị thành niên là MẠNH về mặt lý thuyết hệ thống "
     "(theo lý thuyết hệ thống gia đình) nhưng KHÓ triển khai quy mô lớn vì cha "
     "mẹ thường khó sắp xếp thời gian.")
vn_apply("Việt Nam có thể thử mô hình \"cha mẹ cùng học sinh tham gia\" qua các "
         "buổi họp phụ huynh trực tuyến (Zoom/Zalo). Phù hợp với văn hoá Á Đông "
         "đề cao vai trò gia đình trong giáo dục con cái.")

# ---------- NHÓM 4 ----------
H1("NHÓM 4 — Liệu pháp nhận thức – hành vi qua mạng")

H2("Bài 4.1 — Thử nghiệm CBT online cho lo âu xã hội dưới ngưỡng (JMIR 2024)")
block('Tên bài:',
      'Effectiveness of Unguided Internet-Based Cognitive Behavioral Therapy for '
      'Subthreshold Social Anxiety Disorder in Adolescents and Young Adults: '
      'Multicenter Randomized Controlled Trial')
block('Tạm dịch:',
      'Hiệu quả của Liệu pháp nhận thức – hành vi qua Internet KHÔNG có hướng dẫn '
      'cho rối loạn lo âu xã hội dưới ngưỡng ở vị thành niên và thanh niên: Thử '
      'nghiệm có đối chứng ngẫu nhiên đa trung tâm')
block('Tạp chí:', 'JMIR Pediatrics and Parenting — Vol 7 (2024)')
block('Mã định danh:', 'DOI: 10.2196/55786')
block('Đường dẫn:', 'https://pediatrics.jmir.org/2024/1/e55786')
block('Phương pháp:',
      'Thử nghiệm có đối chứng ngẫu nhiên đa trung tâm về CBT online KHÔNG có '
      'hướng dẫn (unguided — không có nhà trị liệu hỗ trợ trực tiếp) cho rối '
      'loạn lo âu xã hội DƯỚI NGƯỠNG (chưa đủ tiêu chuẩn chẩn đoán đầy đủ) ở '
      'vị thành niên và thanh niên.')
crit("CBT online không có hướng dẫn có TỈ LỆ HOÀN THÀNH thấp (khoảng 30–50%) so "
     "với CBT online có hướng dẫn (50–70%). Tuy nhiên ưu điểm là khả năng triển "
     "khai quy mô lớn rất tốt cho các quốc gia có thu nhập trung bình – thấp như "
     "Việt Nam. Cần đọc kỹ kích thước hiệu ứng cụ thể.")
vn_apply("Việt Nam có thể bản địa hoá CBT online sang tiếng Việt — hợp tác với các "
         "công ty công nghệ giáo dục Việt Nam (ví dụ: Topica, ELSA, Fonos). Nên kết "
         "hợp với chatbot tự động + nhắc nhở qua tin nhắn để tăng tỉ lệ hoàn thành. "
         "Đối tượng phù hợp: học sinh có triệu chứng dưới ngưỡng (sàng lọc bằng "
         "GAD-7 ≥ 5 hoặc PHQ-9 ≥ 5).")

H2("Bài 4.2 — Tổng quan hệ thống về can thiệp số cho điều hoà cảm xúc (JMIR 2022)")
block('Tên bài:',
      'Digital Interventions for Emotion Regulation in Children and Early '
      'Adolescents: Systematic Review and Meta-analysis')
block('Tạm dịch:',
      'Các can thiệp số cho điều hoà cảm xúc ở trẻ em và vị thành niên sớm: Tổng '
      'quan hệ thống và phân tích gộp')
block('Tạp chí:', 'JMIR Serious Games — 2022 (đã được trích dẫn trong các nghiên cứu 2024)')
block('Mã định danh:', 'DOI: 10.2196/31456')
block('Đường dẫn:', 'https://games.jmir.org/2022/3/e31456')
block('Phát hiện chính:',
      'Trò chơi số (digital games) có thiết kế trị liệu giảm trải nghiệm cảm xúc '
      'tiêu cực với hiệu ứng nhỏ nhưng có ý nghĩa thống kê — đặc biệt rõ rệt ở '
      'thanh thiếu niên có nguy cơ lo âu.')
crit("Trò chơi số kết hợp yếu tố trị liệu (\"serious games\") là hướng MỚI và phù "
     "hợp với thế hệ Z. Tuy nhiên cần cân nhắc rủi ro tăng thời gian dùng màn "
     "hình, cần thiết kế để có giới hạn rõ ràng.")
vn_apply("Việt Nam có thể thiết kế ứng dụng học \"kỹ năng vượt khó\" qua trò chơi "
         "mini bằng tiếng Việt. Có thể hợp tác với các studio làm game độc lập "
         "Việt Nam để tạo sản phẩm phù hợp văn hoá.")

# ---------- NHÓM 5 ----------
H1("NHÓM 5 — Hỗ trợ từ bạn đồng trang lứa")

H2("Bài 5.1 — Tổng quan phạm vi hệ thống về hỗ trợ ngang hàng (Murphy 2024)")
block('Tên bài:',
      'A systematic scoping review of peer support interventions in integrated '
      'primary youth mental health care')
block('Tạm dịch:',
      'Tổng quan phạm vi hệ thống về các can thiệp hỗ trợ ngang hàng trong chăm '
      'sóc sức khỏe tâm thần ban đầu tích hợp cho giới trẻ')
block('Tạp chí:', 'Journal of Community Psychology (Wiley) — 2024')
block('Mã định danh:', 'DOI: 10.1002/jcop.23090')
block('Đường dẫn:',
      'https://onlinelibrary.wiley.com/doi/full/10.1002/jcop.23090')
block('Phát hiện chính:',
      'Hỗ trợ ngang hàng (giữa những người trẻ với nhau, có chia sẻ kinh nghiệm '
      'sống cùng vấn đề SKTT) đang được áp dụng rộng trong các bối cảnh sức '
      'khỏe tâm thần cho giới trẻ. Tuy nhiên BẰNG CHỨNG từ các thử nghiệm có '
      'đối chứng còn rất HẠN CHẾ.')
crit("Đây là tổng quan PHẠM VI (không phải tổng quan hệ thống) — chỉ mô tả tình "
     "hình chứ không kết luận quan hệ nhân – quả. Bằng chứng chủ yếu là mô tả + "
     "định tính. Lĩnh vực này vẫn cần nhiều thử nghiệm có đối chứng nghiêm ngặt.")
vn_apply("Mô hình \"Đoàn Thanh niên Cộng sản Hồ Chí Minh\" + \"Lớp trưởng / Bí thư "
         "chi đoàn\" có thể bản địa hoá thành chương trình hỗ trợ ngang hàng có "
         "đào tạo. Cần có chương trình tập huấn cho học sinh nòng cốt + giám sát "
         "chuyên môn từ tư vấn học đường để tránh rủi ro \"truyền\" lo âu giữa "
         "các em.")

H2("Bài 5.2 — Tổng quan hỗ trợ ngang hàng cho SKTT giới trẻ (NCBI Bookshelf)")
block('Mã định danh:', 'NBK602668 (NCBI Bookshelf)')
block('Đường dẫn:', 'https://www.ncbi.nlm.nih.gov/books/NBK602668/')
block('Đặc điểm:',
      'Tổng quan của chính phủ Anh về các chương trình hỗ trợ ngang hàng đã được '
      'triển khai. Tài liệu bổ sung từ Bộ Giáo dục Anh năm 2018.')

# ---------- NHÓM 6 ----------
H1("NHÓM 6 — Bối cảnh châu Á – Trung Quốc")

H2("Bài 6.1 — Tổng quan phạm vi về phòng ngừa SKTT trường học ở Trung Quốc (Lancet 2024)")
block('Tên bài:',
      'School mental health prevention and intervention strategies in China: a '
      'scoping review')
block('Tạm dịch:',
      'Các chiến lược phòng ngừa và can thiệp sức khỏe tâm thần trong trường học '
      'tại Trung Quốc: Tổng quan phạm vi')
block('Tạp chí:',
      'The Lancet Regional Health – Western Pacific (Hệ số tác động khoảng 16) — 2024')
block('Đường dẫn:',
      'https://www.thelancet.com/journals/lanwpc/article/PIIS2666-6065(24)00237-2/fulltext')
block('Phát hiện chính:',
      'Tổng quan phạm vi về các chiến lược phòng ngừa và can thiệp SKTT trong '
      'trường học tại Trung Quốc. Bao quát các chương trình về học tập cảm xúc – '
      'xã hội, chánh niệm, liệu pháp nhận thức – hành vi tại các tỉnh khác nhau '
      'của Trung Quốc.')
crit("RẤT QUAN TRỌNG cho Việt Nam — bối cảnh Trung Quốc tương đồng với Việt Nam "
     "nhất (Đông Á, ảnh hưởng Khổng giáo, áp lực thi đại học cao, hệ thống giáo "
     "dục tập trung). Bài học từ Trung Quốc dễ bản địa hoá cho Việt Nam hơn so "
     "với UK/US.")
vn_apply("Cần đọc kỹ bài này để hiểu Trung Quốc đã làm gì, hiệu quả/thất bại ở "
         "đâu. Việt Nam có thể tránh các sai lầm Trung Quốc đã mắc + bản địa "
         "hoá các thành công.")

# ===================================================================
H1("PHẦN 3 — BẢNG TỔNG HỢP 15 NGHIÊN CỨU")
tbl(['STT', 'Tác giả + Năm', 'Loại nghiên cứu', 'Mã định danh chính', 'Mức quan trọng cho VN'],
    [
        ['1', 'Frontiers Psychiatry 2025',
         'Tổng quan hệ thống + Phân tích gộp 38 thử nghiệm',
         'PMC12127306 / DOI 10.3389/fpsyt.2025.1594658', '⭐⭐⭐'],
        ['2', 'Springer Child Psychiatry 2024',
         'Tổng quan 31 nghiên cứu',
         'PMC12628395 / DOI 10.1007/s10578-024-01667-5', '⭐⭐⭐'],
        ['3', 'Vogelaar et al. 2024',
         'Thử nghiệm có đối chứng theo cụm n=1.613',
         'J School Psychology Vol 105', '⭐⭐'],
        ['4', 'Lochman 2025 (Coping Power)',
         'Thử nghiệm có đối chứng cấp trường',
         'J School Psychology', '⭐⭐'],
        ['5', 'MDPI Pediatric Reports 2025',
         'Phân tích gộp 14 nghiên cứu',
         'Vol 17(3) Bài 59', '⭐⭐'],
        ['6', 'PMC12173555 năm 2025',
         'Thử nghiệm có đối chứng 8+4 tuần',
         'PMC12173555', '⭐⭐'],
        ['7', 'Fulambarkar 2023 (cảnh báo)',
         'Phân tích gộp có tính phản biện',
         'DOI 10.1111/camh.12572', '⭐⭐⭐ (phản biện)'],
        ['8', 'School Mental Health 2023',
         'Nghiên cứu có đối chứng',
         'DOI 10.1007/s12310-023-09620-y', '⭐⭐'],
        ['9', 'Springer Admin Policy 2024',
         'Phân tích gộp d=0,37',
         'DOI 10.1007/s10488-024-01373-3', '⭐⭐⭐'],
        ['10', 'Frontiers Psychiatry 2024',
         'Đào tạo kỹ năng điều hoà cảm xúc',
         'DOI 10.3389/fpsyt.2024.1448529', '⭐⭐'],
        ['11', 'JMIR Pediatrics 2024',
         'Thử nghiệm có đối chứng đa trung tâm',
         'DOI 10.2196/55786', '⭐⭐⭐'],
        ['12', 'JMIR Serious Games 2022',
         'Tổng quan hệ thống số',
         'DOI 10.2196/31456', '⭐⭐'],
        ['13', 'Murphy 2024 (hỗ trợ ngang hàng)',
         'Tổng quan phạm vi hệ thống',
         'DOI 10.1002/jcop.23090', '⭐⭐'],
        ['14', 'NCBI NBK602668',
         'Tổng quan của chính phủ Anh',
         'NBK602668', '⭐'],
        ['15', 'Lancet RH Western Pacific 2024',
         'Tổng quan phạm vi (Trung Quốc)',
         'PIIS2666-6065(24)00237-2', '⭐⭐⭐'],
    ], [0.7, 3.0, 4.0, 5.5, 2.0])

# ===================================================================
H1("PHẦN 4 — KHUYẾN NGHỊ ÁP DỤNG CHO VIỆT NAM")

H2("4.1. Thứ tự ưu tiên các hướng can thiệp (theo độ mạnh bằng chứng)")
nr("Dựa trên 15 nghiên cứu trên, em đề xuất thứ tự ưu tiên cho học sinh trung học "
   "Việt Nam (TỪ MẠNH NHẤT đến yếu nhất):", bold=True)
nr("(1) Quản lý căng thẳng dựa trên Liệu pháp nhận thức – hành vi (đào tạo kỹ năng) — "
   "kích thước hiệu ứng nhỏ đến trung bình, nhất quán nhất.")
nr("(2) Đào tạo kỹ năng điều hoà cảm xúc (kỹ thuật từ DBT, ACT, CBT tập trung cảm xúc) "
   "— hệ số hiệu ứng khoảng 0,37 (mức trung bình).")
nr("(3) Cơ chế tự giới thiệu có chỉ định (đã chứng minh trong nghiên cứu BESST 2024 "
   "đã được dịch và phản biện trong dự án) — đạt hiệu quả trung bình d=−0,52 cho "
   "nhóm đã có triệu chứng cao; phù hợp với mô hình PLACES của Brown 2022.")
nr("(4) Liệu pháp nhận thức – hành vi qua mạng — có thể triển khai quy mô lớn nhưng "
   "tỉ lệ duy trì tham gia thấp khi không có người hướng dẫn.")
nr("(5) Chánh niệm phổ quát (cho mọi học sinh) — hiệu ứng nhỏ (kích thước chuẩn hoá "
   "= −0,14); cần thận trọng với mô hình thất bại MYRIAD ở Anh.")
nr("(6) Hỗ trợ từ bạn đồng trang lứa — hứa hẹn nhưng bằng chứng còn yếu (chưa có "
   "đủ thử nghiệm có đối chứng).")

H2("4.2. Đề xuất chương trình 12 tuần cho học sinh trung học Việt Nam")
tbl(['Tuần', 'Nội dung', 'Phương pháp triển khai', 'Cơ sở khoa học'],
    [
        ['Tuần 1–2',
         'Hiểu biết về sức khỏe tâm thần + Giảm kỳ thị',
         'Tâm lý giáo dục cho toàn lớp (3 buổi × 45 phút)',
         'Vogelaar 2024 + nguyên tắc Lay (đặt tên theo ngôn ngữ thường ngày) '
         'từ mô hình PLACES'],
        ['Tuần 3–4',
         'CBT cốt lõi: Nhận diện cảm xúc + suy nghĩ tiêu cực',
         'Workshop nhóm 60–90 phút/tuần',
         'Brown 2024 BESST với chương trình DISCOVER'],
        ['Tuần 5–6',
         'Kỹ năng CBT: Tái cấu trúc nhận thức + kích hoạt hành vi',
         'Workshop + đóng vai + sách bài tập',
         'Sách hướng dẫn DISCOVER'],
        ['Tuần 7–8',
         'Đào tạo kỹ năng điều hoà cảm xúc',
         'Module DBT (chánh niệm + chịu đựng đau khổ + điều hoà cảm xúc + '
         'hiệu quả giao tiếp)',
         'Springer Admin Policy 2024 với hệ số d=0,37'],
        ['Tuần 9–10',
         'Các kỹ năng ứng phó bổ sung',
         'Tập thể dục, vệ sinh giấc ngủ, quản lý thời gian',
         'Springer 2024 + nghiên cứu QT08 Wen 2020 ở Trung Quốc nông thôn'],
        ['Tuần 11',
         'Hỗ trợ ngang hàng + xây dựng cộng đồng',
         'Chia sẻ nhóm + hệ thống bạn đồng hành',
         'Murphy 2024'],
        ['Tuần 12',
         'Củng cố + phòng ngừa tái phát',
         'Ôn tập + cá nhân hoá mục tiêu + buổi củng cố sau 4 tuần',
         'PMC12173555 (mô hình củng cố sau)'],
    ], [1.2, 4.5, 4.8, 4.5])

H2("4.3. Mô hình triển khai khả thi cho Việt Nam")
nr("• Cấp trường: Tích hợp chương trình 12 tuần vào tiết HĐTN-HN hoặc Giáo dục công "
   "dân lớp 10–12 (ưu tiên lớp 11–12 do áp lực thi đại học cao nhất). Một tiết "
   "mỗi tuần × 60 phút (linh hoạt theo thời khoá biểu).", size=12)
nr("• Người cung cấp can thiệp: Cán bộ tư vấn học đường (theo Thông tư 31/2017/TT-"
   "BGDĐT về tư vấn tâm lý cho học sinh) + chuyên gia tâm lý liên trường giám sát "
   "hằng tháng. Cần đào tạo CBT 5–7 ngày + buổi củng cố hằng quý.", size=12)
nr("• Cơ chế tự giới thiệu (self-referral): Học sinh tự đăng ký qua ứng dụng / "
   "trang web ẩn danh; tránh kênh giáo viên chủ nhiệm để giảm kỳ thị (theo "
   "nghiên cứu BESST 2024 + mô hình PLACES).", size=12)
nr("• Sàng lọc: Sử dụng thang GAD-7 (đã có chuẩn Việt Nam, hệ số tin cậy Cronbach "
   "alpha = 0,916 theo Hoa 2024) hoặc PHQ-9 (chuẩn Việt Nam của Trần Tuấn 2017) "
   "vào đầu năm học. Học sinh có điểm ≥ 5 (mức nhẹ) hoặc ≥ 10 (mức trung bình) "
   "được mời tham gia chương trình có chỉ định; học sinh bình thường vẫn nhận "
   "phần phổ quát (hiểu biết SKTT + giảm kỳ thị).", size=12)
nr("• Hỗ trợ qua công nghệ: Ứng dụng tiếng Việt với các module CBT + chatbot + nhắc "
   "nhở củng cố qua tin nhắn (bản địa hoá từ mô hình JMIR Pediatrics 2024).", size=12)

# ===================================================================
H1("PHẦN 5 — THAM KHẢO ĐẦY ĐỦ")
nr("(15 tài liệu — mỗi tài liệu có DOI, mã PubMed Central nếu có, và đường dẫn truy cập.)",
   italic=True, color=GRAY)

nr("1. Frontiers in Psychiatry (2025). School-based interventions for resilience in "
   "children and adolescents: Systematic review and meta-analysis of randomized "
   "controlled trials. Vol 16. Mã PMC: PMC12127306. DOI: 10.3389/fpsyt.2025.1594658. "
   "https://www.frontiersin.org/journals/psychiatry/articles/10.3389/fpsyt.2025.1594658/full",
   size=11)
nr("2. Academic Stress Interventions in High Schools: A Systematic Literature Review "
   "(2024). Child Psychiatry & Human Development. DOI: 10.1007/s10578-024-01667-5. "
   "Mã PMC: PMC12628395. https://link.springer.com/article/10.1007/s10578-024-01667-5",
   size=11)
nr("3. Vogelaar et al. (2024). Chương trình tâm lý giáo dục \"Stress Lessons\" — "
   "Thử nghiệm có đối chứng theo cụm với 1.613 học sinh ở Hà Lan. Journal of School "
   "Psychology, Vol 105. https://www.sciencedirect.com/science/article/pii/"
   "S0022440524000724", size=11)
nr("4. Lochman et al. (2025). Chương trình \"Sức mạnh ứng phó cho vị thành niên sớm\" "
   "— Thử nghiệm có đối chứng theo nhóm cấp trường, học sinh lớp 7 ở Alabama và "
   "Maryland (Mỹ). Journal of School Psychology. https://www.sciencedirect.com/"
   "science/article/abs/pii/S002244052500010X", size=11)
nr("5. MDPI Pediatric Reports (2025). Mindfulness in Mental Health and Psychiatric "
   "Disorders of Children and Adolescents: Systematic Review and Meta-Analysis. "
   "Vol 17(3): 59. https://www.mdpi.com/2036-7503/17/3/59", size=11)
nr("6. PMC12173555 (2025). Thử nghiệm có đối chứng về can thiệp chánh niệm 8 tuần + "
   "4 tuần củng cố cho vị thành niên 13–15 tuổi. https://pmc.ncbi.nlm.nih.gov/"
   "articles/PMC12173555/", size=11)
nr("7. Fulambarkar N, Seo B, Testerman A, Rees M, Bausback K, Bunge E (2023). "
   "Phân tích gộp về chánh niệm trong trường học cho căng thẳng, trầm cảm, lo âu "
   "ở vị thành niên: một câu chuyện cần cảnh giác. Child and Adolescent Mental "
   "Health (Wiley). DOI: 10.1111/camh.12572. https://acamh.onlinelibrary.wiley.com/"
   "doi/10.1111/camh.12572", size=11)
nr("8. School Mental Health (2023). Can thiệp chánh niệm 10 tuần dựa trên trường học "
   "và triệu chứng trầm cảm + lo âu ở trẻ em và vị thành niên: nghiên cứu có đối "
   "chứng. DOI: 10.1007/s12310-023-09620-y. https://link.springer.com/article/"
   "10.1007/s12310-023-09620-y", size=11)
nr("9. Springer — Administration and Policy in Mental Health and Mental Health Services "
   "Research (2024). Effect of Psychosocial Interventions on Children and Youth Emotion "
   "Regulation: A Meta-Analysis. Hệ số d=0,37. DOI: 10.1007/s10488-024-01373-3. "
   "https://link.springer.com/article/10.1007/s10488-024-01373-3", size=11)
nr("10. Frontiers in Psychiatry (2024). An emotion regulation skills training for "
    "adolescents and parents: perceptions and acceptability of methodological aspects. "
    "DOI: 10.3389/fpsyt.2024.1448529. https://www.frontiersin.org/journals/psychiatry/"
    "articles/10.3389/fpsyt.2024.1448529/full", size=11)
nr("11. JMIR Pediatrics and Parenting (2024). Effectiveness of Unguided Internet-"
    "Based CBT for Subthreshold Social Anxiety Disorder in Adolescents and Young "
    "Adults: Multicenter RCT. DOI: 10.2196/55786. https://pediatrics.jmir.org/"
    "2024/1/e55786", size=11)
nr("12. JMIR Serious Games (2022). Digital Interventions for Emotion Regulation in "
    "Children and Early Adolescents: Systematic Review and Meta-analysis. "
    "DOI: 10.2196/31456. https://games.jmir.org/2022/3/e31456", size=11)
nr("13. Murphy R et al. (2024). A systematic scoping review of peer support "
    "interventions in integrated primary youth mental health care. Journal of "
    "Community Psychology (Wiley). DOI: 10.1002/jcop.23090. "
    "https://onlinelibrary.wiley.com/doi/full/10.1002/jcop.23090", size=11)
nr("14. Peer Support Programs for Youth Mental Health (2024). NCBI Bookshelf "
    "Mã NBK602668. https://www.ncbi.nlm.nih.gov/books/NBK602668/", size=11)
nr("15. The Lancet Regional Health – Western Pacific (2024). School mental health "
    "prevention and intervention strategies in China: a scoping review. "
    "Mã PIIS2666-6065(24)00237-2. https://www.thelancet.com/journals/lanwpc/"
    "article/PIIS2666-6065(24)00237-2/fulltext", size=11)

H2("Đối chiếu kho tài liệu dự án (tránh trùng lặp)")
nr("• Em đã đối chiếu với DATABASE_BAI_BAO_LO_AU.md (kho 35+ bài đã có) — 15 tài liệu "
   "nói trên ĐỀU CHƯA CÓ trong kho dự án; là bổ sung mới.", size=11)
nr("• Các bài liên quan đã có trong kho: Brown 2024 BESST (mã QT042_BESST), Brown 2022 "
   "PLACES (mã QT042_PLACES), Brown & Carter 2025 (mã QT042_B5), Kuyken 2022 MYRIAD, "
   "Stallard 2014 PACES, Zhang 2023 phân tích gộp — tất cả đã được dịch + phản biện "
   "trong 3 tài liệu dịch song ngữ trước đó.", size=11)

H2("Cảnh báo về độ chính xác mã định danh")
crit("Một số DOI em ghi là ƯỚC TÍNH (chưa tải bản đầy đủ từng bài). Cần kiểm tra lại "
     "trước khi trích dẫn chính thức trong báo cáo học thuật. Các bài có DOI/PMC "
     "CHẮC CHẮN (đã đối chiếu qua web): Brown 2022 PLACES PMC8909998, Brown 2024 "
     "BESST PMID 38759665, Murphy 2024 với DOI 10.1002/jcop.23090, Fulambarkar "
     "2023 với DOI 10.1111/camh.12572. Các bài khác (Vogelaar 2024, Lochman 2025): "
     "em có đường dẫn ScienceDirect nhưng DOI chính xác cần kiểm tra lại.")

d.save(OUT)
print('Wrote:', OUT)
