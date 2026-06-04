# -*- coding: utf-8 -*-
"""
RAG v7 — BGE-M3 MULTILINGUAL
Nang cap tu v6 (paraphrase-multilingual-MiniLM-L12-v2) -> v7 (BAAI/bge-m3)

Cai tien ky vong:
- BGE-M3 SOTA cho cross-lingual retrieval (Viet-Anh)
- 1024 chieu (vs 384 cua MiniLM) — grain tinh hon
- 8192 tokens context (vs 512)
- Support dense + sparse + multi-vector retrieval
"""
import os, sys, io, re, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
os.environ['HF_HUB_DISABLE_SYMLINKS_WARNING'] = '1'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
import chromadb
from sentence_transformers import SentenceTransformer

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
DB_PATH = os.path.join(ROOT, 'rag_bao_cao_can_thiep')

# ============================================================
# 1. Collect same sources as v6
# ============================================================
sources = []
report_path = os.path.join(ROOT, '01_Bao-cao', 'Bao cao Can thiep tam ly RLLA VTN - 12042026 v4.docx')
sources.append(('REPORT_v4', report_path, 'report'))

tt_dir = os.path.join(ROOT, 'Tom-tat-tung-bai')
for f in sorted(os.listdir(tt_dir)):
    if f.endswith('.docx') and (f.startswith('VN') or f.startswith('QT')):
        m = re.match(r'(VN\d{3}|QT\d{3})', f)
        if m:
            sources.append((m.group(1), os.path.join(tt_dir, f), 'summary'))

trans_dir = os.path.join(ROOT, '03_Ban-dich')
for f in sorted(os.listdir(trans_dir)):
    if f.endswith('.docx') and (f.startswith('VN') or f.startswith('QT')):
        m = re.match(r'(VN\d{3}|QT\d{3})', f)
        if m:
            sources.append((m.group(1), os.path.join(trans_dir, f), 'translation'))

print(f'Total sources: {len(sources)}')

# ============================================================
# 2. Chunking (same strategy as v6)
# ============================================================
def chunk_doc(path):
    chunks = []
    try:
        doc = Document(path)
    except Exception:
        return chunks

    current_section = ''
    buffer = []

    def flush():
        nonlocal buffer
        if buffer:
            text = '\n'.join(buffer)
            if len(text) > 30:
                chunks.append({'text': text, 'section': current_section})
            buffer = []

    for para in doc.paragraphs:
        t = para.text.strip()
        if not t:
            continue
        style = para.style.name
        if style.startswith('Heading'):
            flush()
            current_section = t
            chunks.append({'text': t, 'section': current_section})
        else:
            buffer.append(t)
    flush()

    for ti, tb in enumerate(doc.tables):
        if len(tb.rows) < 2:
            continue
        headers = [c.text.strip() for c in tb.rows[0].cells]
        parts = ['BẢNG: ' + ' | '.join(headers)]
        for row in tb.rows[1:]:
            cells = [c.text.strip() for c in row.cells]
            pairs = [f'{h}={v}' for h, v in zip(headers, cells) if v]
            if pairs:
                parts.append(' | '.join(pairs))
        chunks.append({'text': '\n'.join(parts), 'section': f'Table {ti+1}'})
    return chunks

all_chunks = []
for cid, path, dtype in sources:
    chunks = chunk_doc(path)
    for i, c in enumerate(chunks):
        all_chunks.append({
            'id': f'{cid}_{dtype}_{i}',
            'text': c['text'],
            'canonical_id': cid,
            'doc_type': dtype,
            'section': c['section'],
            'source_file': os.path.basename(path),
        })

print(f'Total chunks: {len(all_chunks)}')

# ============================================================
# 3. Load BGE-M3 model
# ============================================================
print('\nLoading BAAI/bge-m3...')
model = SentenceTransformer('BAAI/bge-m3', device='cpu')
print(f'Model loaded. Embedding dim: {model.get_sentence_embedding_dimension()}')

# ============================================================
# 4. Build collection
# ============================================================
client = chromadb.PersistentClient(path=DB_PATH)
try:
    client.delete_collection('rag_v7_bgem3')
except Exception:
    pass

col = client.create_collection(
    'rag_v7_bgem3',
    metadata={'hnsw:space': 'cosine',
              'description': 'RAG v7 BGE-M3 multi-source: report v4 + 60 summaries + 60 translations'}
)

texts = [c['text'] for c in all_chunks]
ids = [c['id'] for c in all_chunks]
metas = [{'canonical_id': c['canonical_id'],
          'doc_type': c['doc_type'],
          'section': c['section'][:100],
          'source_file': c['source_file']} for c in all_chunks]

print(f'\nEncoding {len(texts)} chunks with BGE-M3 (1024-dim)...')
batch_size = 16  # smaller batch for BGE-M3 (bigger model)
embs = []
import time
t0 = time.time()
for i in range(0, len(texts), batch_size):
    batch = texts[i:i+batch_size]
    e = model.encode(batch, show_progress_bar=False, normalize_embeddings=True).tolist()
    embs.extend(e)
    if (i // batch_size) % 10 == 0:
        elapsed = time.time() - t0
        rate = (i + len(batch)) / elapsed if elapsed > 0 else 0
        eta = (len(texts) - i - len(batch)) / rate if rate > 0 else 0
        print(f'  {i+len(batch)}/{len(texts)} ({rate:.1f}/s, ETA {eta:.0f}s)')

print(f'  Done in {time.time()-t0:.1f}s')

# Add in batches (Chroma has limit ~5000 per add)
add_batch = 500
for i in range(0, len(ids), add_batch):
    col.add(
        documents=texts[i:i+add_batch],
        embeddings=embs[i:i+add_batch],
        ids=ids[i:i+add_batch],
        metadatas=metas[i:i+add_batch],
    )

print(f'\nAdded {col.count()} chunks to "rag_v7_bgem3"')

# ============================================================
# 5. Smoke tests (same as v6 for comparison)
# ============================================================
print('\n=== SMOKE TEST (10 queries — cùng bộ như v6) ===')

tests = [
    ('Happy House Tran 2023 hiệu quả trầm cảm Cohen d',
     ['Happy House', '0,11', '1.084', 'Tran']),
    ('Dong PLOS 2025 kênh giao tiếp gia đình OR 0,22',
     ['0,22', 'Dong', 'giao tiếp']),
    ('Walder DMHI SAD specific Hedges g 0,878',
     ['0,878', 'Walder', 'SAD']),
    ('Walkup CAMS NEJM đáp ứng CBT + SSRI',
     ['80,7', 'Walkup', 'CAMS']),
    ('Zhu 2025 Suzhou giấc ngủ AOR 13,71',
     ['13,71', 'Zhu', 'giấc ngủ']),
    ('VN030 Happy House RAP-A thích ứng văn hoá',
     ['RAP-A', 'thích ứng', 'Happy House']),
    ('QT040 Walder 21 RCT JMIR',
     ['Walder', '21 RCT', 'JMIR']),
    ('Wen 2020 OR 11,58 rural China',
     ['11,58', 'Wen', 'rural']),
    ('Xian NMA iCBT SUCRA 71,2',
     ['71,2', 'Xian', 'iCBT']),
    ('Tran Thao Vi 2024 Huế longitudinal ESSA',
     ['Thảo Vi', 'Huế', 'ESSA']),
]

passed = 0
for q, expected in tests:
    q_emb = model.encode([q], normalize_embeddings=True).tolist()
    res = col.query(query_embeddings=q_emb, n_results=5,
                    include=['documents', 'distances', 'metadatas'])
    top_text = ' '.join(res['documents'][0]).lower()
    found = [kw for kw in expected if kw.lower() in top_text]
    if len(found) >= 1:
        passed += 1
        status = 'PASS'
    else:
        status = 'FAIL'
    rel = (1 - res['distances'][0][0]) * 100
    top_cid = res['metadatas'][0][0].get('canonical_id', '')
    top_type = res['metadatas'][0][0].get('doc_type', '')
    print(f'[{status}] [{rel:.0f}%] [{top_cid}/{top_type}] "{q[:50]}" found={found}')

print(f'\nSmoke test: {passed}/{len(tests)} PASS ({passed*100//len(tests)}%)')

# Cross-lingual test (Việt query → find English content)
print('\n=== CROSS-LINGUAL TEST (VN query → EN content) ===')
cl_tests = [
    'nghiên cứu can thiệp CBT tại trường học cho lo âu thanh thiếu niên',
    'mối quan hệ giữa giấc ngủ và sức khoẻ tâm thần vị thành niên',
    'so sánh hiệu quả CBT và liệu pháp khác cho lo âu xã hội',
]
for q in cl_tests:
    q_emb = model.encode([q], normalize_embeddings=True).tolist()
    res = col.query(query_embeddings=q_emb, n_results=3,
                    include=['documents', 'metadatas'])
    print(f'\nQ: {q}')
    for i, (doc, meta) in enumerate(zip(res['documents'][0], res['metadatas'][0])):
        cid = meta.get('canonical_id', '')
        print(f'  #{i+1} [{cid}] {doc[:150]}')

# Save manifest
manifest = {
    'collection': 'rag_v7_bgem3',
    'n_chunks': len(all_chunks),
    'n_sources': len(sources),
    'embedding_model': 'BAAI/bge-m3',
    'embedding_dim': 1024,
    'source_report': 'v4',
    'smoke_test_pass_rate': f'{passed}/{len(tests)} ({passed*100//len(tests)}%)',
}
with open(os.path.join(os.path.dirname(__file__), 'rag_v7_manifest.json'), 'w', encoding='utf-8') as f:
    json.dump(manifest, f, ensure_ascii=False, indent=2)
print('\nManifest saved: rag_v7_manifest.json')
