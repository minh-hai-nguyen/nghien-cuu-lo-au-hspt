# -*- coding: utf-8 -*-
"""
verify_all.py — Kiểm tra toàn diện dự án Lo Âu
Chạy: python 06_Scripts/verify_all.py
Idempotent — chạy lại bao nhiêu lần cũng được, không thay đổi dữ liệu.
"""
import os, sys, re, datetime
os.environ['PYTHONIOENCODING'] = 'utf-8'

BASE = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DICH_DIR = os.path.join(BASE, '03_Ban-dich')
TT_DIR = os.path.join(BASE, 'Tom-tat-tung-bai')

errors = []
warnings = []

def check_docx(filepath):
    """Check a DOCX file for common errors"""
    import docx
    doc = docx.Document(filepath)
    text = '\n'.join([p.text for p in doc.paragraphs])
    for t in doc.tables:
        for row in t.rows:
            for cell in row.cells:
                text += ' ' + cell.text
    tables = len(doc.tables)
    return text, tables

print('=' * 60)
print('KIEM TRA TOAN DIEN DU AN LO AU')
print(f'Thoi gian: {datetime.datetime.now().strftime("%Y-%m-%d %H:%M")}')
print('=' * 60)

# ===== 1. Check ban dich =====
print('\n--- 1. BAN DICH ---')
dich_files = sorted([f for f in os.listdir(DICH_DIR)
                     if f.endswith('.docx') and not f.startswith('DICH_')
                     and not f.startswith('~') and not f.startswith('TomTat')])
print(f'  Tong: {len(dich_files)} files')

for f in dich_files:
    try:
        text, tables = check_docx(os.path.join(DICH_DIR, f))

        # Check dot decimal
        bad_dots = re.findall(r'\d+\.\d+%', text)
        if bad_dots:
            errors.append(f'DICH/{f}: DAU CHAM {bad_dots[:3]}')

        # Check phan bien
        if not re.search(r'[Pp]hản biện|PHẢN BIỆN|CRITICAL|phản biện', text, re.IGNORECASE):
            errors.append(f'DICH/{f}: THIEU PHAN BIEN')

        # Check VN diacritics
        vn_chars = len(re.findall(r'[àáảãạăắằẳẵặâấầẩẫậèéẻẽẹêếềểễệìíỉĩịòóỏõọôốồổỗộơớờởỡợùúủũụưứừửữựỳýỷỹỵđ]', text, re.IGNORECASE))
        if vn_chars / max(len(text), 1) < 0.01 and len(text) > 2000:
            errors.append(f'DICH/{f}: MAT DAU TIENG VIET')

    except Exception as e:
        errors.append(f'DICH/{f}: LOI DOC - {str(e)[:50]}')

# ===== 2. Check tom tat =====
print('\n--- 2. TOM TAT ---')
tt_files = sorted([f for f in os.listdir(TT_DIR)
                   if f.endswith('.docx') and (f[0].isdigit() or f.startswith('VN') or f.startswith('QT'))])
print(f'  Tong: {len(tt_files)} files')

for f in tt_files:
    try:
        text, tables = check_docx(os.path.join(TT_DIR, f))

        if tables < 2:
            errors.append(f'TT/{f}: CHI {tables} BANG (can >= 2)')

        if 'Bảng 1' not in text:
            errors.append(f'TT/{f}: THIEU label "Bang 1"')
        if 'Bảng 2' not in text:
            errors.append(f'TT/{f}: THIEU label "Bang 2"')

        bad_dots = re.findall(r'\d+\.\d+%', text)
        if bad_dots:
            errors.append(f'TT/{f}: DAU CHAM {bad_dots[:3]}')

    except Exception as e:
        errors.append(f'TT/{f}: LOI DOC - {str(e)[:50]}')

# ===== 3. Check de cuong =====
print('\n--- 3. DE CUONG + CROSS-STUDY ---')
for fname in os.listdir(BASE):
    if fname.endswith('.docx') and ('cuong' in fname.lower() or 'tong hop' in fname.lower() or 'Tổng hợp' in fname):
        try:
            text, tables = check_docx(os.path.join(BASE, fname))

            # Check key numbers
            key_checks = {
                'Wen 11,58': '11,58' in text,
                'Wen 11,6 (OLD)': not bool(re.search(r'OR\s*=\s*11[,.]6\b', text)),
                'Walkup 2008': 'Walkup' in text if '80,7' in text else True,
            }

            for label, passed in key_checks.items():
                if not passed:
                    errors.append(f'{fname}: LOI TRICH DAN - {label}')

            print(f'  {fname}: {len(text)} chars, {tables} tables')
        except Exception as e:
            warnings.append(f'{fname}: LOI DOC')

# ===== 4. Check 00_Meta/ =====
print('\n--- 4. 00_META ---')
meta_dir = os.path.join(BASE, '00_Meta')
required_files = ['PROJECT_STATE.md', 'PAPERS_INDEX.md', 'CHANGELOG.md', 'ERRORS_FIXED.md']
for f in required_files:
    exists = os.path.exists(os.path.join(meta_dir, f))
    if not exists:
        errors.append(f'00_Meta/{f}: THIEU')
    else:
        print(f'  {f}: OK')

# ===== RESULTS =====
print('\n' + '=' * 60)
if errors:
    print(f'LOI: {len(errors)}')
    for e in errors:
        print(f'  ! {e}')
else:
    print('SACH — 0 LOI')

if warnings:
    print(f'\nCANH BAO: {len(warnings)}')
    for w in warnings:
        print(f'  ~ {w}')

print(f'\nTong: {len(dich_files)} ban dich, {len(tt_files)} tom tat, {len(errors)} loi, {len(warnings)} canh bao')
print('=' * 60)
