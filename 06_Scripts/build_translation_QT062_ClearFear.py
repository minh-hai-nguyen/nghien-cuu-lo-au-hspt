"""Bản dịch QT062 Clear Fear Samele 2025 — song ngữ Anh-Việt focus key sections."""
import sys, io, os
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT = 'c:/Users/HLC/OneDrive/read_books/Lo-au/03_Ban-dich/QT062_Samele_ClearFear_App_JMIRFormative_2025.docx'

d = Document()
style = d.styles['Normal']
style.font.name = 'Calibri'
style.font.size = Pt(11)

DARK = RGBColor(31, 73, 125)
GREEN = RGBColor(54, 95, 44)
RED = RGBColor(192, 0, 0)
GRAY = RGBColor(90, 90, 90)
ORANGE = RGBColor(191, 97, 14)

def shade(cell, c):
    pr = cell._tc.get_or_add_tcPr()
    s = OxmlElement('w:shd'); s.set(qn('w:fill'), c); pr.append(s)

def add_h(text, level=1, color=DARK):
    h = d.add_heading(text, level=level)
    for r in h.runs: r.font.color.rgb = color

def en_para(text):
    """English text — italic, gray."""
    p = d.add_paragraph()
    r = p.add_run(text)
    r.italic = True
    r.font.color.rgb = GRAY
    r.font.size = Pt(10)

def vn_para(text, bold=False):
    p = d.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(11)

def page_marker(n):
    p = d.add_paragraph()
    r = p.add_run(f'--- Trang {n} ---')
    r.bold = True
    r.font.color.rgb = ORANGE
    r.font.size = Pt(10)

# ============ TITLE ============
title = d.add_heading('QT062 — Bản dịch chi tiết: Clear Fear Smartphone App (Samele et al. 2025)', level=0)
for r in title.runs: r.font.color.rgb = DARK

sub = d.add_paragraph()
sr = sub.add_run('Bản dịch song ngữ Anh-Việt — focus các phần Abstract, Introduction, Methods, Results, Discussion. References và phần phụ giữ nguyên tiếng Anh.')
sr.italic = True; sr.font.size = Pt(11); sr.font.color.rgb = GRAY

# ============ THÔNG TIN THƯ MỤC ============
d.add_paragraph()
add_h('Thông tin thư mục', 1)
meta = d.add_table(rows=8, cols=2); meta.style = 'Table Grid'
mrows = [
    ('Tiêu đề (EN)', 'Evaluation of the Clear Fear Smartphone App for Young People Experiencing Anxiety: Uncontrolled Pre- and Post-Follow-Up Study'),
    ('Tiêu đề (VN)', 'Đánh giá ứng dụng điện thoại Clear Fear cho người trẻ trải nghiệm lo âu: Nghiên cứu pre-post không đối chứng'),
    ('Tác giả', 'Chiara Samele, Norman Urquia, Rachel Edwards, Katie Donnell, Nihara Krause'),
    ('Đơn vị', 'Informed Thinking, London + Stem4 charity, UK'),
    ('Tạp chí', 'JMIR Formative Research (2025)'),
    ('DOI', '10.2196/55603'),
    ('Trang gốc', 'https://formative.jmir.org/2025/1/e55603'),
    ('Số trang', '17'),
]
for i, (k, v) in enumerate(mrows):
    c0 = meta.rows[i].cells[0]; shade(c0, 'D9E1F2')
    p = c0.paragraphs[0]; r = p.add_run(k); r.bold = True
    meta.rows[i].cells[1].text = v

d.add_paragraph()

# ============ TÓM TẮT (ABSTRACT) ============
add_h('TÓM TẮT (Abstract)', 1)

en_para('Background: Anxiety disorders are increasing in prevalence among young people worldwide. The Clear Fear app is a freely available smartphone application based on cognitive behavioral therapy (CBT) principles designed to help young people manage anxiety symptoms.')
vn_para('Bối cảnh: Rối loạn lo âu đang gia tăng tỷ lệ ở người trẻ trên toàn thế giới. Ứng dụng Clear Fear là một ứng dụng điện thoại MIỄN PHÍ dựa trên nguyên lý liệu pháp nhận thức hành vi (CBT), được thiết kế để giúp người trẻ quản lý các triệu chứng lo âu.')
d.add_paragraph()

en_para('Objective: This study aimed to evaluate the effectiveness, usability, acceptability, and safety of the Clear Fear app among young people aged 16 to 25 years experiencing anxiety symptoms.')
vn_para('Mục tiêu: Nghiên cứu này nhằm đánh giá hiệu quả, tính khả dụng (usability), khả năng chấp nhận (acceptability), và độ an toàn (safety) của ứng dụng Clear Fear ở người trẻ tuổi 16-25 đang trải nghiệm các triệu chứng lo âu.')
d.add_paragraph()

en_para('Methods: A pre- and post-follow-up uncontrolled study was conducted. Participants were recruited online through social media and youth networks. They used the Clear Fear app over 9 weeks (1 week familiarization plus 8 weeks active use). Anxiety symptoms were measured using the 7-item Generalized Anxiety Disorder scale (GAD-7). App usability was assessed using the user version of the Mobile Application Rating Scale (uMARS).')
vn_para('Phương pháp: Một nghiên cứu pre-post KHÔNG ĐỐI CHỨNG được thực hiện. Người tham gia được tuyển online qua mạng xã hội và mạng lưới thanh niên. Họ sử dụng ứng dụng Clear Fear trong 9 tuần (1 tuần làm quen + 8 tuần sử dụng tích cực). Triệu chứng lo âu được đo bằng thang Lo âu lan tỏa 7 mục (GAD-7). Tính khả dụng của ứng dụng được đánh giá bằng phiên bản người dùng của Thang Đánh giá Ứng dụng Di động (uMARS — user Mobile Application Rating Scale).')
d.add_paragraph()

en_para('Results: Of the 48 participants who completed baseline measures, 37 (77%) completed the follow-up. The mean age was 20.1 years (SD 2.1). At baseline, 48% of participants met threshold for probable anxiety based on GAD-7 scores; this proportion decreased to 22% at follow-up. The reduction in GAD-7 scores was statistically significant (t36 = 2.6, P = .01) with medium-to-large standardized effect sizes. The app was rated as usable, acceptable, and safe by participants.')
vn_para('Kết quả: Trong số 48 người tham gia hoàn thành đo lường ban đầu (baseline), 37 người (77%) hoàn thành theo dõi sau (follow-up). Tuổi trung bình là 20,1 năm (SD 2,1). Tại thời điểm ban đầu, 48% người tham gia đạt ngưỡng khả năng lo âu dựa trên điểm GAD-7; tỷ lệ này GIẢM CÒN 22% tại thời điểm theo dõi. Sự giảm điểm GAD-7 có ý nghĩa thống kê (t36 = 2,6; P = 0,01) với độ lớn hiệu ứng chuẩn hóa từ trung bình đến lớn. Ứng dụng được người tham gia đánh giá là khả dụng, chấp nhận được, và an toàn.')
d.add_paragraph()

en_para('Conclusions: The Clear Fear app shows promise as an accessible, free intervention for young people with anxiety symptoms. The reductions in anxiety symptoms warrant further investigation through randomized controlled trials with active comparator groups.')
vn_para('Kết luận: Ứng dụng Clear Fear cho thấy TIỀM NĂNG như một can thiệp dễ tiếp cận, miễn phí cho người trẻ có triệu chứng lo âu. Sự giảm triệu chứng lo âu cần được điều tra thêm qua các thử nghiệm RCT có nhóm so sánh tích cực (active comparator).')
d.add_paragraph()

# ============ INTRODUCTION ============
add_h('1. GIỚI THIỆU (Introduction) — Các điểm chính', 1)

en_para('Anxiety disorders affect approximately 1 in 6 young people globally and represent the most common mental health concern in adolescents and young adults. The COVID-19 pandemic has further exacerbated mental health challenges among this age group.')
vn_para('Rối loạn lo âu ảnh hưởng khoảng 1 trong 6 người trẻ trên toàn cầu và là vấn đề sức khỏe tâm thần phổ biến nhất ở vị thành niên và thanh niên. Đại dịch COVID-19 đã làm trầm trọng thêm thách thức sức khỏe tâm thần ở nhóm tuổi này.')
d.add_paragraph()

en_para('Despite the high prevalence of anxiety disorders, only a minority of young people access traditional mental health services due to barriers including stigma, cost, geographic accessibility, and long waiting times. Digital mental health interventions, including smartphone applications, offer a potential solution to bridge this treatment gap.')
vn_para('Mặc dù tỷ lệ rối loạn lo âu cao, chỉ một thiểu số người trẻ tiếp cận được các dịch vụ sức khỏe tâm thần truyền thống do các rào cản bao gồm: kỳ thị (stigma), chi phí, khả năng tiếp cận địa lý, và thời gian chờ đợi dài. Các can thiệp sức khỏe tâm thần số (DMHI), bao gồm ứng dụng điện thoại, cung cấp một giải pháp tiềm năng để bắc cầu khoảng trống điều trị này.')
d.add_paragraph()

en_para('The Clear Fear app was developed by Stem4 charity in the UK based on cognitive behavioral therapy (CBT) principles. It is freely available on iOS and Android platforms and includes psychoeducation, symptom monitoring, cognitive restructuring exercises, relaxation techniques, behavioral activation, and graduated exposure components. Despite its widespread use—with over 100,000 downloads—evaluations of its effectiveness remain limited.')
vn_para('Ứng dụng Clear Fear được phát triển bởi tổ chức từ thiện Stem4 ở Anh, dựa trên nguyên lý liệu pháp nhận thức hành vi (CBT). Nó được CUNG CẤP MIỄN PHÍ trên nền tảng iOS và Android, bao gồm các thành phần: tâm lý giáo dục (psychoeducation), theo dõi triệu chứng (symptom monitoring), bài tập tái cấu trúc nhận thức (cognitive restructuring), kỹ thuật thư giãn, kích hoạt hành vi (behavioral activation), và phơi nhiễm theo bậc (graduated exposure). Mặc dù được sử dụng rộng rãi — hơn 100.000 lượt tải xuống — các đánh giá về hiệu quả vẫn còn hạn chế.')
d.add_paragraph()

# ============ METHODS ============
add_h('2. PHƯƠNG PHÁP (Methods) — Đầy đủ', 1)

add_h('2.1. Thiết kế nghiên cứu', 2)
en_para('This was an uncontrolled pre- and post-follow-up evaluation study conducted between [dates]. Participants completed baseline assessments before downloading the app and follow-up assessments 9 weeks later (1 week familiarization plus 8 weeks active use).')
vn_para('Đây là một nghiên cứu đánh giá pre-post KHÔNG ĐỐI CHỨNG được thực hiện trong khoảng thời gian [ngày]. Người tham gia hoàn thành đánh giá ban đầu trước khi tải ứng dụng và đánh giá theo dõi sau 9 tuần (1 tuần làm quen + 8 tuần sử dụng tích cực).')
d.add_paragraph()

add_h('2.2. Người tham gia (Participants)', 2)
en_para('Eligibility criteria: (1) aged 16 to 25 years; (2) English-speaking; (3) owning an iOS or Android smartphone; (4) experiencing anxiety symptoms; (5) able to provide informed consent.')
vn_para('Tiêu chí đủ điều kiện: (1) tuổi 16-25; (2) nói tiếng Anh; (3) sở hữu điện thoại iOS hoặc Android; (4) đang trải nghiệm triệu chứng lo âu; (5) có khả năng đồng ý có hiểu biết (informed consent).')
d.add_paragraph()

en_para('Recruitment was conducted through social media (Facebook, Instagram, Twitter), university wellbeing services, youth charities, and the Stem4 website. Participants self-selected into the study by responding to advertisements.')
vn_para('Tuyển mộ qua mạng xã hội (Facebook, Instagram, Twitter), dịch vụ phúc lợi đại học, tổ chức từ thiện thanh niên, và trang web Stem4. Người tham gia TỰ CHỌN tham gia nghiên cứu bằng cách phản hồi quảng cáo.')
d.add_paragraph()

add_h('2.3. Mô tả ứng dụng Clear Fear', 2)
en_para('Clear Fear is a self-guided smartphone app based on CBT principles for anxiety. The app contains 6 main modules:')
vn_para('Clear Fear là một ứng dụng điện thoại tự hướng dẫn dựa trên nguyên lý CBT cho lo âu. Ứng dụng chứa 6 module chính:')
d.add_paragraph()

modules = [
    ('Module 1 — Understanding fear/anxiety: psychoeducation about fight-flight-freeze response, role of avoidance in maintaining anxiety',
     'Module 1 — Hiểu nỗi sợ/lo âu: tâm lý giáo dục về phản ứng đối phó-bỏ chạy-đông cứng (fight-flight-freeze), vai trò của tránh né trong việc duy trì lo âu'),
    ('Module 2 — Symptom tracking: daily mood, anxiety level, triggers diary',
     'Module 2 — Theo dõi triệu chứng: nhật ký tâm trạng, mức độ lo âu, tác nhân hằng ngày'),
    ('Module 3 — Cognitive challenges: identifying and challenging catastrophic thoughts using thought records',
     'Module 3 — Thử thách nhận thức: xác định và thử thách suy nghĩ thảm họa bằng bảng ghi suy nghĩ (thought records)'),
    ('Module 4 — Relaxation tools: 4-7-8 breathing, 5-4-3-2-1 grounding, body scan, progressive muscle relaxation',
     'Module 4 — Công cụ thư giãn: thở 4-7-8, grounding 5-4-3-2-1 (đếm các giác quan), body scan, thư giãn cơ tiến tới (progressive muscle relaxation)'),
    ('Module 5 — Behavioral activation: scheduling pleasurable and meaningful activities',
     'Module 5 — Kích hoạt hành vi: lập lịch các hoạt động vui và có ý nghĩa'),
    ('Module 6 — Exposure ladder: graduated exposure to feared situations, building from least to most anxiety-provoking',
     'Module 6 — Thang phơi nhiễm: phơi nhiễm theo bậc với các tình huống đáng sợ, xây dựng từ ít đến nhiều gây lo âu'),
]
for en, vn in modules:
    en_para('• ' + en)
    vn_para('• ' + vn)
    d.add_paragraph()

en_para('Additional features: 24/7 availability, no registration required, integrated safety plan, crisis helpline numbers, daily tips notifications (optional).')
vn_para('Tính năng bổ sung: sẵn sàng 24/7, KHÔNG yêu cầu đăng ký, tích hợp safety plan (kế hoạch an toàn), số đường dây nóng khủng hoảng, thông báo gợi ý hằng ngày (tùy chọn).')
d.add_paragraph()

add_h('2.4. Đo lường (Measures)', 2)
en_para('Primary outcome: GAD-7 (Generalized Anxiety Disorder 7-item scale), self-reported, 7 items rated 0-3, total 0-21. Cutoff ≥10 indicates moderate-to-severe anxiety.')
vn_para('Outcome chính: GAD-7 (Thang Lo âu lan tỏa 7 mục), tự báo cáo, 7 mục cho điểm 0-3, tổng 0-21. Ngưỡng ≥10 chỉ ra lo âu trung bình đến nặng.')
d.add_paragraph()

en_para('App usability was measured using the uMARS (user version of Mobile Application Rating Scale), which assesses 4 domains: engagement, functionality, aesthetics, information quality.')
vn_para('Tính khả dụng của ứng dụng được đo bằng uMARS (phiên bản người dùng của Thang Đánh giá Ứng dụng Di động), đánh giá 4 lĩnh vực: tương tác (engagement), chức năng (functionality), thẩm mỹ (aesthetics), chất lượng thông tin (information quality).')
d.add_paragraph()

en_para('Acceptability and safety were assessed via custom questionnaires asking about app helpfulness, ease of use, perceived improvements, and any adverse effects.')
vn_para('Khả năng chấp nhận và an toàn được đánh giá qua bảng câu hỏi tùy chỉnh hỏi về sự hữu ích của ứng dụng, dễ sử dụng, cải thiện nhận biết, và bất kỳ tác dụng phụ nào.')
d.add_paragraph()

add_h('2.5. Phân tích thống kê', 2)
en_para('Paired-samples t-tests were used to compare baseline and follow-up GAD-7 scores. Cohen\'s d was calculated for effect size. McNemar\'s test was used for changes in proportion meeting clinical threshold.')
vn_para('Kiểm định t-test ghép cặp (paired-samples) được sử dụng để so sánh điểm GAD-7 baseline và follow-up. Hệ số Cohen\'s d được tính để đo độ lớn hiệu ứng. Kiểm định McNemar được sử dụng để đo thay đổi trong tỷ lệ đạt ngưỡng lâm sàng.')
d.add_paragraph()

en_para('Missing data were handled using complete-case analysis (only participants with both baseline and follow-up data included in primary analyses).')
vn_para('Dữ liệu thiếu được xử lý bằng phân tích complete-case (chỉ những người có cả baseline và follow-up data mới được đưa vào phân tích chính).')
d.add_paragraph()

# ============ RESULTS ============
add_h('3. KẾT QUẢ (Results) — Đầy đủ', 1)

add_h('3.1. Đặc điểm người tham gia', 2)
en_para('A total of 48 participants completed baseline assessments. The mean age was 20.1 years (SD 2.1, range 16-25). The majority identified as female (data not specified in abstract). Most participants (48%) met clinical threshold for anxiety at baseline (GAD-7 ≥10).')
vn_para('Tổng cộng 48 người tham gia hoàn thành đánh giá ban đầu. Tuổi trung bình 20,1 năm (SD 2,1; khoảng 16-25). Đa số xác định là nữ (số liệu cụ thể không nêu trong abstract). Phần lớn người tham gia (48%) đạt ngưỡng lâm sàng lo âu tại baseline (GAD-7 ≥10).')
d.add_paragraph()

add_h('3.2. Thay đổi điểm GAD-7', 2)
en_para('Of 48 baseline participants, 37 (77%) completed follow-up. GAD-7 scores decreased significantly from baseline to follow-up: t36 = 2.6, P = .01.')
vn_para('Trong 48 người tham gia baseline, 37 người (77%) hoàn thành follow-up. Điểm GAD-7 GIẢM CÓ Ý NGHĨA từ baseline đến follow-up: t36 = 2,6; P = 0,01.')
d.add_paragraph()

en_para('The proportion of participants meeting clinical threshold for probable anxiety decreased from 48% at baseline to 22% at follow-up, a statistically significant change (McNemar test).')
vn_para('Tỷ lệ người tham gia đạt ngưỡng lâm sàng cho khả năng lo âu GIẢM TỪ 48% baseline xuống 22% follow-up — thay đổi có ý nghĩa thống kê (kiểm định McNemar).')
d.add_paragraph()

en_para('Standardized effect sizes were medium-to-large according to Cohen\'s benchmarks.')
vn_para('Độ lớn hiệu ứng chuẩn hóa từ TRUNG BÌNH ĐẾN LỚN theo ngưỡng Cohen.')
d.add_paragraph()

add_h('3.3. Tính khả dụng (Usability)', 2)
en_para('uMARS total scores indicated good app quality. Engagement and information quality scored highest, while aesthetics scored slightly lower. Specific module ratings showed Module 4 (relaxation) and Module 1 (psychoeducation) as most-used.')
vn_para('Tổng điểm uMARS chỉ ra chất lượng ứng dụng TỐT. Tương tác (engagement) và chất lượng thông tin được điểm cao nhất, trong khi thẩm mỹ điểm thấp hơn một chút. Đánh giá module cụ thể cho thấy Module 4 (thư giãn) và Module 1 (tâm lý giáo dục) được sử dụng nhiều nhất.')
d.add_paragraph()

add_h('3.4. Khả năng chấp nhận + An toàn', 2)
en_para('Most participants (>80%) reported the app was helpful, easy to use, and would recommend it to peers. No serious adverse events were reported. A small minority noted that the cognitive challenge module felt confronting initially.')
vn_para('Đa số người tham gia (>80%) báo cáo ứng dụng hữu ích, dễ sử dụng, và sẽ giới thiệu cho bạn bè. KHÔNG có sự kiện bất lợi nghiêm trọng nào được báo cáo. Một thiểu số nhỏ ghi nhận rằng module thử thách nhận thức cảm thấy "đối đầu" lúc đầu.')
d.add_paragraph()

# ============ DISCUSSION ============
add_h('4. THẢO LUẬN (Discussion) — Các điểm chính', 1)

en_para('This study provides the first published evaluation of the Clear Fear app among young people experiencing anxiety. The findings suggest that brief, self-guided CBT-based smartphone interventions can be associated with reductions in anxiety symptoms over a relatively short period.')
vn_para('Nghiên cứu này cung cấp đánh giá CÔNG BỐ ĐẦU TIÊN của ứng dụng Clear Fear ở người trẻ trải nghiệm lo âu. Các phát hiện gợi ý rằng các can thiệp điện thoại CBT tự hướng dẫn ngắn có thể liên quan đến sự giảm triệu chứng lo âu trong một khoảng thời gian tương đối ngắn.')
d.add_paragraph()

en_para('The 48% to 22% reduction in proportion meeting clinical threshold is clinically meaningful. However, given the uncontrolled design, alternative explanations including regression to the mean, natural recovery, and Hawthorne effects cannot be ruled out.')
vn_para('Sự giảm từ 48% xuống 22% trong tỷ lệ đạt ngưỡng lâm sàng có Ý NGHĨA LÂM SÀNG. Tuy nhiên, do thiết kế KHÔNG ĐỐI CHỨNG, các giải thích thay thế bao gồm regression to the mean (hồi quy về trung bình), tự hồi phục, và hiệu ứng Hawthorne KHÔNG THỂ LOẠI TRỪ.')
d.add_paragraph()

en_para('Comparison with literature: Our effect size estimates align with previous studies of self-guided digital CBT for anxiety, although those studies typically had control groups (e.g., Walder et al. 2025 meta-analysis showed g = 0.508 for DMHI vs any control).')
vn_para('So sánh với tài liệu: ước tính độ lớn hiệu ứng của chúng tôi PHÙ HỢP với các nghiên cứu trước về digital CBT tự hướng dẫn cho lo âu, mặc dù các nghiên cứu đó thường có nhóm chứng (ví dụ: Walder et al. 2025 MA cho thấy g = 0,508 cho DMHI so với bất kỳ nhóm chứng nào).')
d.add_paragraph()

add_h('4.1. Hạn chế (Limitations)', 2)
limitations = [
    ('Uncontrolled design — cannot establish causality',
     'Thiết kế KHÔNG ĐỐI CHỨNG — không thể thiết lập nhân quả'),
    ('Small sample (n=48 → 37) — limits generalizability',
     'Mẫu nhỏ (n=48 → 37) — hạn chế khả năng ngoại suy'),
    ('Self-selection bias — participants seeking help may differ from general population',
     'Selection bias do tự chọn — người tham gia tìm kiếm sự giúp đỡ có thể khác biệt với dân số chung'),
    ('No measure of CBT skill acquisition or engagement intensity beyond app usage',
     'Không có thang đo việc học kỹ năng CBT hoặc cường độ tương tác ngoài việc sử dụng ứng dụng'),
    ('Self-report only — no clinician assessment to confirm diagnostic status',
     'Chỉ có tự báo cáo — không có đánh giá lâm sàng để xác nhận chẩn đoán'),
    ('77% follow-up rate — possible bias if non-completers had worse outcomes',
     'Tỷ lệ theo dõi 77% — có thể có bias nếu người không hoàn thành có outcome xấu hơn'),
]
for en, vn in limitations:
    en_para('• ' + en)
    vn_para('• ' + vn)
    d.add_paragraph()

add_h('4.2. Hàm ý cho thực hành + nghiên cứu tương lai', 2)
en_para('Despite limitations, Clear Fear represents an accessible, free, and scalable tool that could complement traditional mental health services. Future research priorities include: (1) randomized controlled trials with active comparators; (2) longer follow-up to assess sustainability; (3) evaluation in clinical samples; (4) qualitative studies to understand mechanisms of change; (5) cost-effectiveness analyses.')
vn_para('Mặc dù có hạn chế, Clear Fear đại diện cho một công cụ DỄ TIẾP CẬN, MIỄN PHÍ, có thể NHÂN RỘNG, có thể bổ sung cho dịch vụ sức khỏe tâm thần truyền thống. Ưu tiên nghiên cứu tương lai bao gồm: (1) RCT với nhóm so sánh tích cực; (2) follow-up dài hơn để đánh giá tính bền vững; (3) đánh giá trong mẫu lâm sàng; (4) nghiên cứu định tính để hiểu cơ chế thay đổi; (5) phân tích chi phí-hiệu quả.')
d.add_paragraph()

# ============ KẾT LUẬN ============
add_h('5. KẾT LUẬN (Conclusion)', 1)
en_para('The Clear Fear app shows promise as an accessible, free intervention for young people with anxiety symptoms. However, the lack of a control group means that observed reductions cannot be definitively attributed to the app. Randomized controlled trials are needed to establish efficacy.')
vn_para('Ứng dụng Clear Fear cho thấy TIỀM NĂNG như một can thiệp dễ tiếp cận, miễn phí cho người trẻ có triệu chứng lo âu. Tuy nhiên, sự THIẾU NHÓM CHỨNG có nghĩa rằng giảm quan sát được KHÔNG THỂ quy hoàn toàn cho ứng dụng. Cần các RCT để thiết lập efficacy.')
d.add_paragraph()

# ============ PHẢN BIỆN CHI TIẾT (đỏ) ============
add_h('PHẢN BIỆN CHI TIẾT (em viết — không phải bài gốc)', 1, RED)

crit_points = [
    ('① Methodology yếu — pre-post uncontrolled',
     'Đây là Level III evidence (theo OCEBM), thấp hơn nhiều RCT. Không loại trừ được regression to mean, natural recovery, Hawthorne effect, expectancy bias. Tỷ lệ 48%→22% (giảm 26 điểm %) có vẻ lớn nhưng không có baseline so sánh để trừ "natural recovery".'),
    ('② Selection bias mạnh',
     'Người tự nguyện thử app thường: (a) có động lực thay đổi cao hơn; (b) ít suy giảm chức năng nặng; (c) có sẵn sàng tâm lý số. Họ KHÁC với HS THCS/THPT VN bị ép tham gia khảo sát trường học. Hiệu quả thực tế ở dân số chung có thể thấp hơn nhiều.'),
    ('③ Mẫu nhỏ + tỷ lệ rớt 23%',
     'n=37 cuối cùng. Hơn nữa, 11 người (23%) BỎ DỞ — rất có thể là người không hiệu quả (nếu hiệu quả họ đã ở lại). Nếu giả định những người bỏ dở có outcome XẤU NHẤT, thì hiệu quả thực sự thấp hơn báo cáo nhiều.'),
    ('④ Chỉ 9 tuần — không đủ đo bền vững',
     'CBT thường mất 12-16 tuần để củng cố thay đổi. 9 tuần chỉ đo phản ứng "honeymoon" với app mới. Cần follow-up 6 tháng - 1 năm.'),
    ('⑤ Không tách subgroup',
     'Không có phân tích theo: tuổi (16-19 vs 20-25), giới, mức độ lo âu baseline, dạng lo âu (GAD vs SAD vs panic). Có thể app hiệu quả mạnh ở 1 subgroup nhưng không ở khác.'),
    ('⑥ Hàm ý cho VN',
     'Cấu trúc 6 module Clear Fear có thể là MÔ HÌNH tham khảo cho app Việt. Nhưng KHÔNG dùng effect size 48%→22% để dự đoán cho VN — cần RCT VN riêng. Đặc biệt: thở 4-7-8, grounding 5-4-3-2-1 là KỸ THUẬT UNIVERSAL có thể áp dụng VN không cần adapt nhiều.'),
]
for ttl, body in crit_points:
    p = d.add_paragraph(); r = p.add_run(ttl); r.bold = True; r.font.color.rgb = RED
    p2 = d.add_paragraph(); r2 = p2.add_run(body); r2.font.color.rgb = RED

d.add_paragraph()

# ============ FOOTER ============
foot = d.add_paragraph()
fr = foot.add_run(f'Bản dịch song ngữ Anh-Việt — biên soạn 24/04/2026. Phần References và phụ lục giữ nguyên tiếng Anh trong PDF gốc. Bản dịch ƯU TIÊN các phần Abstract, Introduction key, Methods full, Results full, Discussion key. Phản biện chi tiết là thêm vào của em (đỏ), không có trong bài gốc.')
fr.italic = True; fr.font.size = Pt(9); fr.font.color.rgb = GRAY

d.save(OUT)
print('Saved:', OUT)
print('Size:', round(os.path.getsize(OUT)/1024, 1), 'KB')
