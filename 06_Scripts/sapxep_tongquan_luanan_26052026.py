# -*- coding: utf-8 -*-
"""Sap xep lai phan Tong quan nghien cuu cua luan an TS theo nhom chu de.
SUA NHE: giu nguyen 100% noi dung khoa hoc, chi thay doi heading va bo cac
sub-heading "Nghien cuu o nuoc ngoai / Viet Nam" de gom VN vao cuoi tung chu de.
Them muc Yeu to bao ve voi placeholder [Can bo sung].
Ngay 26/05/2026."""
import os, sys, io, json
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

import re
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_COLOR_INDEX
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
SRC = os.path.join(ROOT, 'Luận án TS', '01_Rối loạn lo âu_ 26-05.docx')
OUT = os.path.join(ROOT, 'Luận án TS', '01_TongQuan_NhomTheoChuDe_v1_26052026.docx')

# ============================================================
# DOC SETUP
# ============================================================
def doc_init():
    doc = Document()
    s = doc.styles['Normal']
    s.font.name = 'Times New Roman'; s.font.size = Pt(13)
    s.paragraph_format.space_after = Pt(6); s.paragraph_format.line_spacing = 1.5
    for sec in doc.sections:
        sec.top_margin = Cm(3.0); sec.bottom_margin = Cm(2.5)
        sec.left_margin = Cm(3.5); sec.right_margin = Cm(2.5)
    return doc

def H(doc, text, level=2):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.name = 'Times New Roman'; r.font.color.rgb = RGBColor(0,0,0)
    return h

def P(doc, text, fli_cm=1.25, italic=False, color=None):
    """Tao 1 paragraph noi dung. fli_cm = first_line_indent cm (None/0 = khong indent)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if fli_cm is not None and fli_cm > 0:
        p.paragraph_format.first_line_indent = Cm(fli_cm)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    if italic:
        r.italic = True
    if color is not None:
        r.font.color.rgb = color
    return p

def P_bullet(doc, text):
    """List bullet style for paragraphs originally List Bullet."""
    p = doc.add_paragraph(style='List Bullet')
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    return p

def P_placeholder(doc, text):
    """[Can bo sung] placeholder - in nghieng + chu DO DAM (khong highlight de de edit)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Cm(1.25)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    r.italic = True
    r.bold = True
    r.font.color.rgb = RGBColor(192, 0, 0)
    return p

def add_meta_paragraph(doc, text, fli_cm=0):
    """Doan metadata cua file - in nghieng + chu DO DAM (khong highlight)."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if fli_cm > 0:
        p.paragraph_format.first_line_indent = Cm(fli_cm)
    r = p.add_run(text)
    r.font.name = 'Times New Roman'; r.font.size = Pt(13)
    r.italic = True
    r.bold = True
    r.font.color.rgb = RGBColor(192, 0, 0)
    return p


# ============================================================
# DOC CO COPY NGUYEN VAN TU LUAN AN GOC
# Doc truc tiep tu file goc de lay duoc paragraph_format (FLI, etc.)
# ============================================================
SRC_DOC = Document(SRC)
SRC_PARAS = SRC_DOC.paragraphs

def copy_paragraph_from_src(out_doc, src_idx):
    """Copy paragraph tu file goc voi format chinh xac:
    - First-line indent
    - Line spacing
    - Run-level: bold, italic, underline, font name, font size
    - Tab/space prefix giu nguyen
    """
    if src_idx >= len(SRC_PARAS):
        return None
    sp = SRC_PARAS[src_idx]
    if not sp.text.strip():
        return None

    style = sp.style.name
    src_fli = sp.paragraph_format.first_line_indent
    fli_cm = src_fli.cm if src_fli is not None else 0
    src_li = sp.paragraph_format.left_indent
    li_cm = src_li.cm if src_li is not None else 0
    src_ri = sp.paragraph_format.right_indent
    ri_cm = src_ri.cm if src_ri is not None else 0
    src_ls = sp.paragraph_format.line_spacing
    src_sa = sp.paragraph_format.space_after
    src_align = sp.alignment

    # Tao paragraph moi
    if style == 'List Bullet':
        new_p = out_doc.add_paragraph(style='List Bullet')
    elif style in ('List Continue 3', 'macro'):
        new_p = out_doc.add_paragraph(style='List Bullet')
    else:
        new_p = out_doc.add_paragraph()
        # Alignment - giu y nguyen, default JUSTIFY neu None
        new_p.alignment = src_align if src_align is not None else WD_ALIGN_PARAGRAPH.JUSTIFY
        if fli_cm and fli_cm > 0:
            new_p.paragraph_format.first_line_indent = Cm(fli_cm)
        if li_cm and li_cm > 0:
            new_p.paragraph_format.left_indent = Cm(li_cm)
        if ri_cm and ri_cm > 0:
            new_p.paragraph_format.right_indent = Cm(ri_cm)
        if src_ls is not None:
            new_p.paragraph_format.line_spacing = src_ls
        if src_sa is not None:
            new_p.paragraph_format.space_after = src_sa

    # Copy tung run voi format (bold, italic, underline, font)
    for src_run in sp.runs:
        if not src_run.text:
            continue
        new_run = new_p.add_run(src_run.text)
        new_run.bold = src_run.bold
        new_run.italic = src_run.italic
        new_run.underline = src_run.underline
        # Font
        font_name = src_run.font.name or 'Times New Roman'
        new_run.font.name = font_name
        font_size = src_run.font.size
        new_run.font.size = font_size if font_size else Pt(13)
        # Color (chi copy neu khac default)
        if src_run.font.color and src_run.font.color.rgb:
            try:
                new_run.font.color.rgb = src_run.font.color.rgb
            except Exception:
                pass
    return new_p


# ============================================================
# LOAD CONTENT TU JSON DA TRICH XUAT (chi dung cho idx_map de reference)
# ============================================================
with open(os.path.join(ROOT, '06_Scripts', '_workspace', 'tongquan_raw.json'),
          'r', encoding='utf-8') as f:
    items = json.load(f)

idx_map = {it['idx']: it for it in items}

# Danh sach paragraph cu bi loai bo (la heading cu ma se duoc thay bang heading moi)
SKIP_IDX = {
    218,  # H2 "Tong quan nghien cuu" - se them lai
    219,  # "Nghien cuu xac dinh ty le pho bien..." - se thanh H2 moi
    221,  # "Nghien cuu o nuoc ngoai"
    259,  # "Nghien cuu o Viet Nam"
    310,  # "Nhan xet ve khoang trong nghien cuu" - se thanh H3 moi
    321,  # "Nghien cuu ve cac yeu to anh huong..." - se thanh H2 moi
    322,  # "Nghien cuu o nuoc ngoai"
    323,  # "Yeu to truong hoc" - se thanh H3 moi
    350,  # "Yeu to hoc sinh" - se thanh H3 moi
    372,  # "Yeu to gia dinh" - se thanh H3 moi
    384,  # "Nghien cuu o Viet Nam ve cac yeu to..."
    385,  # "Yeu to truong hoc" (VN, trung lap)
    392,  # "Yeu to hoc sinh" (VN, trung lap)
    403,  # "Yeu to gia dinh" (VN, trung lap)
    410,  # "Nhan xet ve khoang trong nghien cuu cac yeu to..." - se thanh H3
    413,  # "Nghien cuu ve can thiep..." - se thanh H2 moi
    415,  # "Nghien cuu o nuoc ngoai"
    482,  # "Nghien cuu o Viet Nam ve can thiep..."
    487,  # "Nhan xet ve khoang trong nghien cuu ve can thiep..." - se thanh H3
}

def emit_range(doc, start, end, skip=SKIP_IDX):
    """Emit cac paragraph trong khoang [start, end] inclusive, bo qua SKIP_IDX.
    Doc truc tiep tu file goc de giu nguyen paragraph_format (FLI)."""
    n = 0
    for i in range(start, end + 1):
        if i in skip:
            continue
        if i not in idx_map:
            continue
        if copy_paragraph_from_src(doc, i) is not None:
            n += 1
    return n


# ============================================================
# BUILD DOC MOI
# ============================================================
doc = doc_init()

# Mo dau (heading H1 + metadata in highlight de NCS de NHAN BIET phai EDIT)
H(doc, 'TỔNG QUAN NGHIÊN CỨU', 1)
add_meta_paragraph(doc, '(Bản sắp xếp lại theo nhóm chủ đề — 26/05/2026 — XÓA dòng này sau khi đã chèn vào luận án)')
add_meta_paragraph(doc, 'GHI CHÚ CHO NCS: Phiên bản này sắp xếp lại các nghiên cứu trong Tổng quan của luận án theo nhóm chủ đề, gộp các nghiên cứu Việt Nam vào cuối mỗi nhóm thay vì để ở mục riêng. Toàn bộ nội dung khoa học, tài liệu trích dẫn và số liệu được giữ nguyên 100% từ bản gốc; chỉ có heading và một số câu chuyển ý ngắn được điều chỉnh. Mục về Yếu tố bảo vệ là phần thầy hướng dẫn yêu cầu bổ sung — hiện đang để placeholder [Cần bổ sung] vì Tổng quan hiện tại chưa có nội dung tương ứng. CÁC PHẦN TÔ ĐỎ + VÀNG (như đoạn này) là chỗ NCS cần xử lý/xóa/thay thế bằng tay.')
doc.add_paragraph()

# Doan intro mo bai (para 220)
emit_range(doc, 220, 220)


# ============================================================
# MUC 1: NGHIEN CUU VE TY LE, MUC DO, BIEU HIEN
# ============================================================
H(doc, '1. Nghiên cứu về tỷ lệ phổ biến, mức độ và biểu hiện của rối loạn lo âu ở trẻ vị thành niên', 2)
# QT + VN trong cung 1 dong chay (skip 221 "Nghien cuu o nuoc ngoai" va 259 "Nghien cuu o Viet Nam")
emit_range(doc, 222, 309)

# Nhan xet, khoang trong VN
H(doc, '1.1. Nhận xét và khoảng trống ở Việt Nam', 3)
emit_range(doc, 311, 320)


# ============================================================
# MUC 2: NGHIEN CUU VE YEU TO NGUY CO
# ============================================================
H(doc, '2. Nghiên cứu về yếu tố nguy cơ rối loạn lo âu', 2)

# 2.1 Yeu to truong hoc - QT (324-349) + VN (386-391)
H(doc, '2.1. Yếu tố trường học', 3)
emit_range(doc, 324, 349)
emit_range(doc, 386, 391)

# 2.2 Yeu to hoc sinh - QT (351-371) + VN (393-402)
H(doc, '2.2. Yếu tố học sinh', 3)
emit_range(doc, 351, 371)
emit_range(doc, 393, 402)

# 2.3 Yeu to gia dinh - QT (373-383) + VN (404-409)
H(doc, '2.3. Yếu tố gia đình', 3)
emit_range(doc, 373, 383)
emit_range(doc, 404, 409)

# 2.4 Nhan xet, khoang trong VN
H(doc, '2.4. Nhận xét và khoảng trống ở Việt Nam', 3)
emit_range(doc, 411, 412)


# ============================================================
# MUC 3: NGHIEN CUU VE YEU TO BAO VE (MOI - PLACEHOLDER)
# ============================================================
H(doc, '3. Nghiên cứu về yếu tố bảo vệ', 2)
add_meta_paragraph(doc, '[Phần này được thêm theo yêu cầu của thầy hướng dẫn. Hiện Tổng quan hiện tại của luận án chưa có nội dung đầy đủ về các yếu tố bảo vệ. NCS cần bổ sung theo các tiểu mục sau, dựa trên tài liệu thực tế đã thu thập, KHÔNG bịa số liệu hay tài liệu mới.]')

H(doc, '3.1. Lòng tự trọng', 3)
P_placeholder(doc, '[Cần bổ sung — Tổng quan hiện tại chưa có nội dung về chủ đề này. NCS có thể bắt đầu từ các tài liệu về thang đo Rosenberg và phân tích tổng hợp của Sowislo và Orth (2013) đã trích dẫn trong Bài 1 — Yếu tố nguy cơ.]')

H(doc, '3.2. Gắn bó trường học', 3)
P_placeholder(doc, '[Cần bổ sung — Tổng quan hiện tại chưa có nội dung về chủ đề này. Hướng tham khảo: school connectedness, school engagement.]')

H(doc, '3.3. Hỗ trợ của cha mẹ', 3)
P_placeholder(doc, '[Cần bổ sung — Tổng quan hiện tại chưa có nội dung về chủ đề này. Hướng tham khảo: parental warmth, parental support.]')

H(doc, '3.4. Hỗ trợ từ bạn bè', 3)
P_placeholder(doc, '[Cần bổ sung — Tổng quan hiện tại chưa có nội dung về chủ đề này. Hướng tham khảo: peer support, peer attachment.]')

H(doc, '3.5. Biện pháp ứng phó của học sinh', 3)
P_placeholder(doc, '[Cần bổ sung — Tổng quan hiện tại chưa có nội dung về chủ đề này. NCS có thể bắt đầu từ phân tích tổng hợp của Compas và cộng sự (2017) đã trích dẫn trong Bài 1 — Yếu tố nguy cơ, về phân biệt coping problem-focused / emotion-focused và approach / avoidance.]')

H(doc, '3.6. Nhận xét và khoảng trống ở Việt Nam', 3)
P_placeholder(doc, '[Cần bổ sung sau khi đã có nội dung 5 tiểu mục trên.]')


# ============================================================
# MUC 4: NGHIEN CUU VE CAN THIEP
# ============================================================
H(doc, '4. Nghiên cứu về can thiệp rối loạn lo âu cho trẻ vị thành niên', 2)
# intro can thiep (414)
emit_range(doc, 414, 414)
# QT can thiep (416-481), VN can thiep (483-486), bo qua 415 "Nghien cuu o nuoc ngoai" va 482 "Nghien cuu o VN"
emit_range(doc, 416, 481)
emit_range(doc, 483, 486)

# Nhan xet, khoang trong VN
H(doc, '4.1. Nhận xét và khoảng trống ở Việt Nam', 3)
emit_range(doc, 488, 488)


# ============================================================
# LUU
# ============================================================
doc.save(OUT)

# Bao cao
from docx import Document as D
d_out = D(OUT)
words = sum(len(p.text.split()) for p in d_out.paragraphs)
n_para = sum(1 for p in d_out.paragraphs if p.text.strip())
size = os.path.getsize(OUT) // 1024
print(f"Da luu: {OUT}")
print(f"Tu: {words} | Doan: {n_para} | Size: {size}KB")

# So sanh voi file goc
total_orig = sum(len(it['text'].split()) for it in items)
print(f"\nTu goc (218-489): {total_orig}")
print(f"Tu moi: {words}")
print(f"Chenh lech: {words - total_orig:+d} tu ({(words - total_orig) / total_orig * 100:+.1f}%)")