# -*- coding: utf-8 -*-
"""Batch dịch QT27-34"""
import sys, os
sys.path.insert(0, os.path.dirname(os.path.abspath(__file__)))
from tao_dich_template import *

OUT_DIR = os.path.dirname(os.path.abspath(__file__))

def save(doc, name):
    path = os.path.join(OUT_DIR, name)
    doc.save(path)
    import docx as dx
    d2 = dx.Document(path)
    t = '\n'.join([p.text for p in d2.paragraphs])
    print(f'  -> {name}: {len(t)} chars, {len(d2.tables)} tables')

# =====================================================================
# QT27 — Nature Human Behaviour 2025 — Fassi et al. — Social Media
# =====================================================================
def dich_QT27():
    print('QT27 Nature SocialMedia...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1038/s41562-025-02134-4', size=10)
    add_heading(doc, 'Sử dụng mạng xã hội ở thanh thiếu niên có và không có tình trạng sức khỏe tâm thần', 1)
    h = doc.add_paragraph(); r = h.add_run('Social Media Use in Adolescents with and without Mental Health Conditions'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Social Media Use in Adolescents with and without Mental Health Conditions'),
        ('Tác giả', 'Luisa Fassi, Amanda M. Ferguson, Andrew K. Przybylski, Tamsin J. Ford, Amy Orben'),
        ('Cơ quan', 'MRC Cognition and Brain Sciences Unit, Cambridge; Oxford Internet Institute; Đại học Cambridge'),
        ('Tạp chí', 'Nature Human Behaviour (Q1, IF ≈ 30)'),
        ('Xuất bản', '2025, Vol. 9, pp. 1283–1299, 21 trang'),
        ('DOI', '10.1038/s41562-025-02134-4'),
        ('Loại NC', 'Registered Report — cắt ngang, dữ liệu MHCYP (Mental Health of Children and Young People)'),
        ('Mẫu', 'N = 3.340 VTN Anh (11–19 tuổi), đại diện quốc gia, chẩn đoán lâm sàng'),
    ])
    add_page_ref(doc, '1283–1299', 'Nature Human Behaviour', 'Vol. 9, 2025')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Lo ngại về mối quan hệ giữa sử dụng MXH và SKTT VTN ngày càng tăng, nhưng ít NC tập trung vào VTN có triệu chứng SKTT ở mức lâm sàng. Registered Report này phân tích dữ liệu đại diện quốc gia UK (N = 3.340, 11–19 tuổi) bao gồm đánh giá chẩn đoán bởi nhà lâm sàng cùng với đo lường MXH định lượng và định tính.')

    p = doc.add_paragraph()
    r = p.add_run('Kết quả: VTN có tình trạng SKTT dành NHIỀU thời gian hơn trên MXH (g = 0,46; KTC 90%: 0,38–0,54) — khác biệt có ý nghĩa và có thể quan trọng. Tuy nhiên, so sánh xã hội trực tuyến (g = 0,30), thiếu kiểm soát thời gian (g = 0,27), và tác động phản hồi (g = 0,29) — mặc dù có ý nghĩa thống kê nhưng NẰM TRONG NGƯỠNG TƯƠNG ĐƯƠNG (SESOI g = 0,4). Nghĩa là: khác biệt tồn tại nhưng CÓ THỂ NHỎ. VTN có rối loạn NỘI HÓA (lo âu, trầm cảm) bị ảnh hưởng NHIỀU HƠN VTN có rối loạn NGOẠI HÓA.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['Nature Human Behaviour Q1 IF ≈ 30 — tạp chí uy tín CỰC CAO.','Registered Report — phương pháp tiền đăng ký, giảm thiên lệch xuất bản.','Dữ liệu MHCYP — chẩn đoán LÂM SÀNG (không chỉ sàng lọc), đại diện quốc gia UK.','Phân biệt nội hóa vs ngoại hóa — rất quan trọng cho lý thuyết.','Đo cả định lượng (thời gian) VÀ định tính (so sánh, kiểm soát, phản hồi).']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Cắt ngang — không xác lập nhân quả.','Chỉ UK — văn hóa MXH có thể khác châu Á/VN.','16% có SKTT — mẫu lâm sàng nhỏ.','Tự báo cáo thời gian MXH.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. GIỚI THIỆU', 2)
    add_p(doc, 'Lo ngại về MXH và SKTT VTN ngày càng tăng, nhưng hầu hết NC chỉ dùng sàng lọc (không chẩn đoán lâm sàng). Điều này hạn chế hiểu biết về mối quan hệ ở VTN thực sự có rối loạn. Registered Report này sử dụng dữ liệu MHCYP (Mental Health of Children and Young People in England, NHS Digital) — khảo sát quốc gia với chẩn đoán bởi nhà lâm sàng dựa trên SDQ và DAWBA.')
    add_p(doc, 'Phân tích cả MXH định lượng (thời gian) và định tính (so sánh xã hội, kiểm soát, phản hồi, cyberbullying). So sánh VTN có/không có tình trạng SKTT, phân biệt nội hóa (lo âu, trầm cảm) vs ngoại hóa (ADHD, rối loạn hành vi).')

    add_heading(doc, '2. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Dữ liệu: MHCYP 2017, NHS Digital. N = 3.340 VTN 11–19 tuổi, đại diện quốc gia Anh. Chẩn đoán SKTT bởi nhà lâm sàng dựa trên DAWBA (Development and Well-Being Assessment). 50% nam, 50% nữ. 16% có ≥1 tình trạng SKTT (N = 519); 8% nội hóa (N = 282); 3% ngoại hóa (N = 104).')
    add_p(doc, 'Phân tích: Registered Report — phương pháp tiền đăng ký. So sánh nhóm có/không SKTT trên nhiều chiều MXH. Hedges\' g, NHST + equivalence testing (SESOI g = 0,4).')

    add_heading(doc, '3. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. So sánh sử dụng MXH: VTN có vs không có SKTT', 3)
    add_table(doc,
        ['Chiều MXH', 'Hedges\' g', 'KTC 90%', 'NHST p', 'Tương đương?', 'Kết luận'],
        [['Thời gian trên MXH', '0,46', '0,38–0,54', '<0,001', 'Không (>SESOI)', 'Khác biệt CÓ Ý NGHĨA'],
         ['So sánh xã hội trực tuyến', '0,30', '0,22–0,39', '<0,001', 'Có (<SESOI)', 'Có nhưng NHỎ'],
         ['Thiếu kiểm soát thời gian', '0,27', '0,19–0,35', '<0,001', 'Có', 'Có nhưng NHỎ'],
         ['Tác động phản hồi', '0,29', '0,21–0,38', '<0,001', 'Có', 'Có nhưng NHỎ'],
         ['Bắt nạt mạng', '—', '—', '—', '—', 'Xem bài']],
        widths=[3.5, 2.0, 2.0, 1.5, 2.5, 3.0])
    add_p(doc, 'SESOI (Smallest Effect Size of Interest) = g = 0,4. "Tương đương" = khác biệt < ngưỡng có ý nghĩa thực tế.', size=9, italic=True)

    add_heading(doc, 'Bảng 2. So sánh nội hóa vs ngoại hóa', 3)
    add_table(doc,
        ['Nhóm', 'Thời gian MXH (g)', 'So sánh XH (g)', 'Phản hồi (g)', 'Ghi chú'],
        [['Nội hóa (lo âu, trầm cảm)', '0,47', '0,42', '0,43', 'Ảnh hưởng NHIỀU hơn'],
         ['Ngoại hóa (ADHD, hành vi)', '0,38', '0,12', '0,08', 'Ảnh hưởng ÍT hơn']],
        widths=[4.0, 3.0, 3.0, 2.5, 3.5])
    add_p(doc, 'VTN rối loạn nội hóa nhạy cảm hơn với MXH — phù hợp lý thuyết (rumination, so sánh tiêu cực).', size=9, italic=True)

    add_heading(doc, '4. THẢO LUẬN VÀ KẾT LUẬN', 2)
    add_p(doc, 'VTN có tình trạng SKTT dùng MXH nhiều hơn và trải nghiệm MXH tiêu cực hơn — nhưng hầu hết khác biệt NHỎ (dưới SESOI). Ngoại lệ: thời gian trên MXH (g = 0,46 > SESOI 0,4). VTN nội hóa bị ảnh hưởng nhiều hơn ngoại hóa. Kết quả không ủng hộ mạnh mẽ giả thuyết MXH GÂY HẠI lớn, nhưng cũng không bác bỏ — cần NC dọc và thực nghiệm.')

    add_abbreviation_table(doc, [('MXH','Mạng xã hội'),('SKTT','Sức khỏe tâm thần'),('MHCYP','Mental Health of Children and Young People — SKTT Trẻ em và Thanh niên'),('DAWBA','Development and Well-Being Assessment'),('SDQ','Strengths and Difficulties Questionnaire'),('SESOI','Smallest Effect Size of Interest — Kích thước Hiệu ứng Nhỏ nhất Đáng quan tâm'),('VTN','Vị thành niên')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['Nature Human Behaviour Q1 IF ≈ 30 — cực uy tín. Registered Report.','Chẩn đoán lâm sàng (DAWBA) — vượt trội so với sàng lọc.','Equivalence testing — phương pháp tiên tiến, phân biệt "có ý nghĩa" vs "đủ lớn".','Phân biệt nội hóa vs ngoại hóa — đóng góp lý thuyết quan trọng.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['Cắt ngang — không xác lập nhân quả (MXH→SKTT hay SKTT→MXH?).','Chỉ UK 2017 — dữ liệu trước TikTok bùng nổ. MXH 2024 khác 2017.','16% có SKTT — mẫu lâm sàng vẫn nhỏ cho phân tích tiểu nhóm.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN chưa có NC MXH với chẩn đoán lâm sàng — đa số dùng sàng lọc.','Cần NC nội hóa vs ngoại hóa ở VN — đặc biệt lo âu + MXH.','So sánh tác động MXH ở VN vs UK — văn hóa khác.']:
        add_red(doc, f'• {s}')
    save(doc, '27_NatureHumanBehav_SocialMedia_2025.docx')

# =====================================================================
# QT28 — AJP 2024 — Zugman et al. — Pediatric Anxiety Treatment
# =====================================================================
def dich_QT28():
    print('QT28 AJP Treatment...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1176/appi.ajp.20231037', size=10)
    add_heading(doc, 'Phương pháp tiếp cận hiện tại và tương lai trong điều trị rối loạn lo âu ở trẻ em', 1)
    h = doc.add_paragraph(); r = h.add_run('Current and Future Approaches to Pediatric Anxiety Disorder Treatment'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Current and Future Approaches to Pediatric Anxiety Disorder Treatment'),
        ('Tác giả', 'Andre Zugman, MD, PhD; Anderson M. Winkler, MD, PhD; Purnima Qamar, BS; Daniel S. Pine, MD'),
        ('Cơ quan', 'NIMH (National Institute of Mental Health), NIH, Bethesda, MD, Hoa Kỳ'),
        ('Tạp chí', 'American Journal of Psychiatry — AJP (Q1, IF ≈ 18)'),
        ('Xuất bản', '2024, Vol. 181, pp. 189–200, 12 trang'),
        ('DOI', '10.1176/appi.ajp.20231037'),
        ('Loại tài liệu', 'Tổng quan (Overview) — đánh giá phê bình tài liệu'),
    ])
    add_page_ref(doc, '189–200', 'AJP', 'Vol. 181, 2024')

    add_heading(doc, 'TÓM TẮT', 2)
    add_p(doc, 'Tổng quan đánh giá phê bình tài liệu về điều trị rối loạn lo âu trẻ em. Hai phương pháp điều trị đã được thiết lập: liệu pháp nhận thức–hành vi (CBT — Cognitive-Behavioral Therapy) và thuốc chống trầm cảm (SSRIs). Nhiều trẻ nhận điều trị không đạt thuyên giảm → cần phương pháp mới. Bài viết tóm tắt CBT, thuốc hiện có, và NC thần kinh khoa học đặt nền tảng cho cải thiện.')

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['AJP Q1 IF ≈ 18 — tạp chí tâm thần hàng đầu thế giới.','Tác giả từ NIMH/NIH — cơ quan NC SKTT hàng đầu Mỹ.','Tổng quan toàn diện CBT + SSRIs + hướng tương lai.','Kết hợp CBT+SSRI hiệu quả nhất: CAMS 80,7% đáp ứng.','Đề xuất hướng mới: attention bias modification, D-cycloserine, thực tế ảo.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. ĐIỀU TRỊ ĐÃ THIẾT LẬP', 2)
    add_p(doc, '1.1. Dược lý', bold=True)
    add_p(doc, 'SSRIs (ức chế tái hấp thu serotonin chọn lọc) là điều trị dược lý hàng đầu. Phân tích tổng hợp cung cấp bằng chứng mạnh mẽ về hiệu quả: NNT (Number Needed to Treat) = 4 cho đáp ứng. Fluoxetine, sertraline, fluvoxamine — hiệu quả tương đương. SNRIs (duloxetine, venlafaxine) cũng có hiệu quả nhưng ít dữ liệu hơn.')
    add_p(doc, 'Lưu ý an toàn: FDA cảnh báo hộp đen về nguy cơ tự tử với tất cả SSRIs/SNRIs ở VTN. Tuy nhiên, phân tích tổng hợp cho thấy nguy cơ thấp và lợi ích vượt trội rủi ro.')
    add_p(doc, 'Thuốc thay thế: tricyclics (kém hiệu quả, nhiều tác dụng phụ), benzodiazepines (chỉ 3 RCT nhỏ, không khuyến nghị), buspirone (ít dữ liệu).')

    add_p(doc, '1.2. CBT', bold=True)
    add_p(doc, 'CBT là điều trị tâm lý hàng đầu. Phân tích tổng hợp cho thấy hiệu quả trung bình–lớn (NNT ≈ 3–4). Các thành phần chính: giáo dục tâm lý, tái cấu trúc nhận thức, phơi nhiễm dần (exposure therapy). Hiệu quả duy trì 6–12 tháng sau điều trị.')

    add_p(doc, '1.3. Kết hợp CBT + SSRI', bold=True)
    add_p(doc, 'CAMS (Child/Adolescent Anxiety Multimodal Study) — RCT lớn nhất: kết hợp CBT+sertraline đạt 80,7% đáp ứng vs CBT đơn 59,7% vs sertraline đơn 54,9% vs placebo 23,7%. Kết hợp VƯỢT TRỘI cả đơn trị.')

    add_heading(doc, 'Bảng 1. Hiệu quả điều trị rối loạn lo âu trẻ em', 3)
    add_table(doc,
        ['Phương pháp', 'Tỷ lệ đáp ứng', 'NNT', 'Ghi chú'],
        [['CBT + SSRI (kết hợp)', '80,7%', '~2', 'CAMS — hiệu quả nhất'],
         ['CBT đơn', '59,7%', '~3', 'Phơi nhiễm là thành phần chính'],
         ['SSRI đơn (sertraline)', '54,9%', '~4', 'Cảnh báo FDA về tự tử'],
         ['Placebo', '23,7%', '—', 'Hiệu ứng placebo đáng kể']],
        widths=[4.0, 3.0, 2.0, 5.0])

    add_heading(doc, '2. HƯỚNG TƯƠNG LAI', 2)
    add_p(doc, '• Attention Bias Modification Training (ABMT): huấn luyện chú ý tránh xa kích thích đe dọa. Kết quả lẫn lộn nhưng hứa hẹn khi kết hợp với CBT.')
    add_p(doc, '• D-cycloserine (DCS): thuốc tăng cường học tập phơi nhiễm. Ít bằng chứng ở trẻ em.')
    add_p(doc, '• Thực tế ảo (VR): phơi nhiễm qua VR — thuận tiện, kiểm soát, nhưng ít RCT.')
    add_p(doc, '• Biomarker: fMRI phản ứng amygdala, EEG — dự báo đáp ứng điều trị.')

    add_heading(doc, '3. KẾT LUẬN', 2)
    add_p(doc, 'CBT và SSRIs là điều trị đã thiết lập cho rối loạn lo âu trẻ em, với kết hợp hiệu quả nhất (80,7%). Tuy nhiên, ~20–40% không đáp ứng — cần phương pháp mới. NC thần kinh khoa học (amygdala, attention bias) đặt nền tảng cho cải thiện, nhưng cần nhiều RCT hơn.')

    add_abbreviation_table(doc, [('CBT','Cognitive-Behavioral Therapy — Liệu pháp Nhận thức–Hành vi'),('SSRI','Selective Serotonin Reuptake Inhibitor — Thuốc Ức chế Tái hấp thu Serotonin Chọn lọc'),('SNRI','Serotonin-Norepinephrine Reuptake Inhibitor'),('CAMS','Child/Adolescent Anxiety Multimodal Study'),('NNT','Number Needed to Treat — Số Bệnh nhân Cần Điều trị'),('ABMT','Attention Bias Modification Training'),('VR','Virtual Reality — Thực tế Ảo'),('RCT','Randomized Controlled Trial'),('NIMH','National Institute of Mental Health'),('FDA','Food and Drug Administration')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['AJP Q1 IF ≈ 18. Tác giả NIMH — chuyên gia hàng đầu thế giới.','Tổng quan toàn diện: CBT + dược lý + kết hợp + hướng tương lai.','CAMS: RCT lớn nhất, kết hợp CBT+SSRI 80,7% đáp ứng.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['Tổng quan narrative — không phải phân tích tổng hợp hệ thống.','Chủ yếu dữ liệu phương Tây — ít đề cập châu Á/LMIC.','Hướng tương lai (ABMT, DCS, VR) chưa có đủ bằng chứng.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN chưa có RCT CBT cho lo âu trẻ em — GAP rất lớn.','Cần đánh giá hiệu quả CBT trong bối cảnh VN (văn hóa, nguồn lực).','Kết hợp CBT+SSRI chưa được thử nghiệm tại VN.']:
        add_red(doc, f'• {s}')
    save(doc, '28_AJP_PediatricAnxiety_2024.docx')

# =====================================================================
# QT29 — BMC Psychiatry 2025 — Li et al. — CBT Network Meta-analysis
# =====================================================================
def dich_QT29():
    print('QT29 CBT NetworkMeta...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1186/s12888-025-07227-y', size=10)
    add_heading(doc, 'Hiệu quả các loại can thiệp khác nhau cho rối loạn lo âu ở trẻ em và thanh thiếu niên: Tổng quan hệ thống và phân tích tổng hợp mạng', 1)
    h = doc.add_paragraph(); r = h.add_run('Effectiveness of Different Types of Interventions for Anxiety Disorders in Children and Adolescents: A Systematic Review and Network Meta-Analysis'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Effectiveness of Different Types of Interventions for Anxiety Disorders in Children and Adolescents: A Systematic Review and Network Meta-Analysis'),
        ('Tác giả', 'Li et al.'),
        ('Cơ quan', 'Capital University of Physical Education and Sports, Beijing; Shenyang Pharmaceutical University'),
        ('Tạp chí', 'BMC Psychiatry (Q1, IF ≈ 4,4)'),
        ('Xuất bản', '2025, 25:809, 14 trang'),
        ('DOI', '10.1186/s12888-025-07227-y'),
        ('Loại NC', 'Tổng quan hệ thống + phân tích tổng hợp mạng (network meta-analysis)'),
    ])
    add_page_ref(doc, '1–14', 'BMC Psychiatry', '2025, 25:809')

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Bối cảnh: Rối loạn lo âu phổ biến nhất ở trẻ em/VTN, tăng sau COVID. NC nhằm đánh giá hiệu quả các can thiệp khác nhau và xác định phương pháp hiệu quả nhất. Phương pháp: Tìm kiếm hệ thống RCT từ PubMed, EMBASE, Cochrane, PsycINFO. Phân tích tổng hợp mạng (NMA) so sánh nhiều can thiệp đồng thời. Kết quả: CBT cá nhân và CBT nhóm đều hiệu quả đáng kể so với đối chứng. CBT cá nhân xếp hạng cao nhất. Hoạt động thể chất cũng cho thấy hiệu quả.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['BMC Psychiatry Q1. Network meta-analysis — so sánh nhiều can thiệp.','CBT cá nhân xếp hạng cao nhất.','Hoạt động thể chất cũng hiệu quả — phát hiện thú vị.','Sau COVID-19 — bối cảnh thời sự.']:
        add_p(doc, f'• {b}')

    add_heading(doc, '1. PHƯƠNG PHÁP', 2)
    add_p(doc, 'Tìm kiếm hệ thống trong PubMed, EMBASE, Cochrane Library, PsycINFO. Tiêu chuẩn: RCT, trẻ em/VTN với lo âu, so sánh ≥2 can thiệp. Phân tích tổng hợp mạng (NMA) sử dụng mô hình Bayesian, xếp hạng bằng SUCRA (Surface Under the Cumulative Ranking).')

    add_heading(doc, '2. KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xếp hạng hiệu quả can thiệp cho lo âu trẻ em/VTN', 3)
    add_table(doc,
        ['Can thiệp', 'SUCRA (%)', 'Xếp hạng', 'So với đối chứng'],
        [['CBT cá nhân', 'Cao nhất', '1', 'Hiệu quả đáng kể'],
         ['CBT nhóm', 'Cao', '2', 'Hiệu quả đáng kể'],
         ['CBT kết hợp (cá nhân+nhóm)', 'Cao', '3', 'Hiệu quả'],
         ['Hoạt động thể chất', 'Trung bình–Cao', '4–5', 'Hiệu quả — đáng chú ý'],
         ['Thuốc (SSRIs)', 'Trung bình', '—', 'Hiệu quả (không so sánh trực tiếp NMA)'],
         ['Danh sách chờ/Placebo', '—', 'Thấp nhất', 'Reference']],
        widths=[4.0, 2.5, 2.0, 5.0])

    add_heading(doc, '3. KẾT LUẬN', 2)
    add_p(doc, 'CBT cá nhân là can thiệp hiệu quả nhất cho rối loạn lo âu trẻ em/VTN, tiếp theo là CBT nhóm. Hoạt động thể chất cũng hiệu quả — phù hợp làm can thiệp bổ sung tại trường. Kết quả ủng hộ CBT là tiêu chuẩn vàng, phù hợp với AJP 2024 (Zugman et al.).')

    add_abbreviation_table(doc, [('NMA','Network Meta-Analysis — Phân tích Tổng hợp Mạng'),('SUCRA','Surface Under the Cumulative Ranking — Diện tích Dưới Xếp hạng Tích lũy'),('CBT','Cognitive-Behavioral Therapy'),('RCT','Randomized Controlled Trial'),('SSRI','Selective Serotonin Reuptake Inhibitor')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['NMA — phương pháp mạnh, so sánh gián tiếp nhiều can thiệp.','Bao gồm hoạt động thể chất — ít NMA nào làm.','Kết quả nhất quán với tổng quan trước (CBT hàng đầu).']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['Hầu hết RCT từ phương Tây — ít dữ liệu châu Á.','Không phân biệt loại lo âu (GAD vs social anxiety vs phobia).','Chất lượng RCT khác nhau — một số mẫu nhỏ.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN chưa có RCT CBT cho lo âu VTN.','Hoạt động thể chất tại trường VN — có thể triển khai dễ hơn CBT. Cần NC.','CBT nhóm phù hợp bối cảnh VN (ít chuyên gia, nhiều VTN) — cần thử nghiệm.']:
        add_red(doc, f'• {s}')
    save(doc, '29_CBT_NetworkMeta_2025_BMCPsych.docx')

# =====================================================================
# QT30 — GBD Trends 2025 — Zhang et al. — J Affective Disorders
# =====================================================================
def dich_QT30():
    print('QT30 GBD Trends...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.04.001', size=10)
    add_heading(doc, 'Xu hướng rối loạn trầm cảm và lo âu ở VTN và thanh niên (10–24 tuổi) từ 1990 đến 2021: Phân tích Gánh nặng Bệnh tật Toàn cầu', 1)
    h = doc.add_paragraph(); r = h.add_run('Trends in Depressive and Anxiety Disorders among Adolescents and Young Adults (aged 10–24) from 1990 to 2021: A Global Burden of Disease Study Analysis'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tiêu đề gốc', 'Trends in Depressive and Anxiety Disorders among Adolescents and Young Adults (aged 10–24) from 1990 to 2021'),
        ('Tác giả', 'Zhang Dongjun, Wu Mingyue, Li Xinqi, Wang Lina, Wu Jiali, Jin Mengyao'),
        ('Cơ quan', 'Đại học Y Xinxiang, Trung Quốc'),
        ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
        ('Xuất bản', '2025, 14 trang'),
        ('DOI', '10.1016/j.jad.2025.04.001'),
        ('Loại NC', 'Phân tích xu hướng sinh thái hồi cứu — dữ liệu GBD 2021'),
        ('Phạm vi', '204 quốc gia, 1990–2021, VTN 10–24 tuổi'),
    ])

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: AAPC toàn cầu cho trầm cảm và lo âu ở 10–24 tuổi: 0,80%–0,97%/năm (1990–2021), tăng tốc từ 2014–2021, đỉnh 2019. Gánh nặng cao nhất ở nữ. Tỷ lệ mới mắc trầm cảm tăng từ 3.085,5/100.000 (1990) lên 3.909,9/100.000 (2021), AAPC = 0,97%. Lo âu: 708,0 → 883,1/100.000, AAPC = 0,84%. Bất bình đẳng: AAPC cao nhất ở nước SDI cao, thấp nhất ở SDI thấp.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'Bảng 1. Xu hướng toàn cầu trầm cảm và lo âu ở 10–24 tuổi', 3)
    add_table(doc,
        ['Rối loạn', 'Tỷ lệ mới mắc 1990 (/100K)', 'Tỷ lệ mới mắc 2021 (/100K)', 'AAPC', 'p'],
        [['Trầm cảm', '3.085,5', '3.909,9', '0,97%', '<0,001'],
         ['Lo âu', '708,0', '883,1', '0,84%', '<0,001'],
         ['Trầm cảm (nam)', '2.265,2', '2.991,7', '1,07%', '<0,001'],
         ['Trầm cảm (nữ)', '3.932,7', '4.875,1', '0,93%', '<0,001'],
         ['Lo âu (10–14)', '397,1', '547,2', '1,12%', '<0,001'],
         ['Lo âu (15–19)', '943,6', '1.180,8', '0,93%', '<0,001'],
         ['Lo âu (20–24)', '1.105,8', '1.287,1', '0,60%', '<0,001']],
        widths=[3.5, 3.5, 3.5, 2.0, 1.5])
    add_p(doc, 'AAPC = Average Annual Percent Change. GBD 2021, 204 quốc gia. UI = Uncertainty Interval.', size=9, italic=True)

    add_heading(doc, 'Bảng 2. Xu hướng theo khu vực (AAPC cao nhất)', 3)
    add_table(doc,
        ['Khu vực', 'AAPC Trầm cảm', 'AAPC Lo âu', 'Ghi chú'],
        [['Bắc Mỹ thu nhập cao', '2,30%', '—', 'Cao nhất'],
         ['Mỹ Latinh Trung', '2,00%', '1,48%', ''],
         ['Mỹ Latinh Andes', '1,68%', '1,91%', 'Lo âu cao nhất'],
         ['Đông Á', 'Trung bình', 'Trung bình', 'Bao gồm VN'],
         ['SDI cao', 'Cao nhất', 'Cao nhất', 'Tương quan SDI'],
         ['SDI thấp', 'Thấp nhất', 'Thấp nhất', 'Có thể do thiếu dữ liệu']],
        widths=[4.0, 3.0, 3.0, 4.0])

    add_heading(doc, 'KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu GBD 204 quốc gia 31 năm cho thấy trầm cảm và lo âu ở 10–24 tuổi tăng liên tục toàn cầu (AAPC 0,84–0,97%), tăng tốc từ 2014, đỉnh 2019. Nữ gánh nặng cao hơn. Nước SDI cao tăng nhanh nhất — bất bình đẳng giữa các nước.')

    add_abbreviation_table(doc, [('GBD','Global Burden of Disease — Gánh nặng Bệnh tật Toàn cầu'),('AAPC','Average Annual Percent Change — Thay đổi Phần trăm Trung bình Hàng năm'),('DALYs','Disability-Adjusted Life Years — Năm Sống Điều chỉnh Tàn tật'),('SDI','Socio-Demographic Index — Chỉ số Kinh tế–Nhân khẩu Xã hội'),('UI','Uncertainty Interval — Khoảng Bất định')])

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh:', bold=True)
    for s in ['GBD 2021 — dữ liệu toàn diện nhất, 204 quốc gia, 31 năm.','Joinpoint regression — xác định điểm thay đổi xu hướng (2014 tăng tốc).','Phân tích bất bình đẳng (SII, CI) — ít NC nào làm.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Hạn chế:', bold=True)
    for s in ['GBD dựa trên mô hình ước tính — nhiều nước (bao gồm VN) thiếu dữ liệu gốc.','Không phân biệt loại lo âu (GAD, SAD, phobia).','Sinh thái hồi cứu — không xác lập nhân quả.']:
        add_red(doc, f'• {s}')
    add_red(doc, 'Gap:', bold=True)
    for s in ['VN CHƯA CÓ dữ liệu xu hướng dài hạn lo âu VTN — GBD ước tính có thể không chính xác.','Cần khảo sát quốc gia lặp lại (như MHCYP UK) để xác nhận xu hướng VN.']:
        add_red(doc, f'• {s}')
    save(doc, '30_GBD_Trends_10-24y_2025.docx')

# =====================================================================
# QT31 — 59 Countries — Islam et al. — J Affective Disorders
# =====================================================================
def dich_QT31():
    print('QT31 59Countries...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1016/j.jad.2025.06.001', size=10)
    add_heading(doc, 'Tỷ lệ và các yếu tố liên quan của lo âu ở thanh thiếu niên đi học: Phân tích từ 59 quốc gia', 1)
    h = doc.add_paragraph(); r = h.add_run('Prevalence of and Factors Associated with Anxiety among School Going Adolescents: Analysis from 59 Countries'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Md. Amirul Islam, Tanjirul Islam, Bristi Rani Saha, et al.'),
        ('Cơ quan', 'Xi\'an Jiao Tong University; Khulna University, Bangladesh'),
        ('Tạp chí', 'Journal of Affective Disorders (Q1, IF ≈ 6,6)'),
        ('Xuất bản', '2025, 11 trang'),
        ('Loại NC', 'Phân tích dữ liệu thứ cấp — GSHS (Global School-based Health Survey)'),
        ('Mẫu', '59 quốc gia, VTN đi học 13–17 tuổi'),
    ])

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Phân tích dữ liệu GSHS (WHO) từ 59 quốc gia. Tỷ lệ lo âu chung ở VTN đi học. Yếu tố liên quan: bất an thực phẩm (food insecurity), bắt nạt, cô đơn, thiếu ngủ. Phân tích đa biến xác định yếu tố nguy cơ và bảo vệ.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'KẾT QUẢ CHÍNH', 2)
    add_p(doc, '• 59 quốc gia LMIC (Low and Middle-Income Countries) — đại diện GSHS.')
    add_p(doc, '• Tỷ lệ lo âu: khác nhau theo khu vực, cao nhất ở Trung Đông và châu Phi.')
    add_p(doc, '• Yếu tố nguy cơ mạnh nhất: bất an thực phẩm, cô đơn, bắt nạt, thiếu ngủ.')
    add_p(doc, '• Yếu tố bảo vệ: hỗ trợ từ bạn bè, cha mẹ, hoạt động thể chất.')

    add_heading(doc, 'KẾT LUẬN', 2)
    add_p(doc, 'Lo âu ở VTN đi học phổ biến ở 59 quốc gia LMIC, liên quan mạnh với bất an thực phẩm và bắt nạt. Cần can thiệp đa tầng: an ninh lương thực, chống bắt nạt, tăng hỗ trợ xã hội.')

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: 59 quốc gia, GSHS đại diện. Bao gồm LMIC — phù hợp VN.', bold=True)
    add_red(doc, 'Hạn chế: Cắt ngang. GSHS chỉ VTN đi học (loại bỏ thất nghiệp, bỏ học).')
    add_red(doc, 'Gap: VN có dữ liệu GSHS nhưng chưa phân tích riêng lo âu theo bối cảnh này.')
    save(doc, '31_59Countries_Anxiety_2025.docx')

# =====================================================================
# QT32 — Ireland My World — Fitzgerald et al. 2024
# =====================================================================
def dich_QT32():
    print('QT32 Ireland...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1017/eip.2024.1', size=10)
    add_heading(doc, 'Khám phá xu hướng thay đổi trầm cảm và lo âu ở thanh thiếu niên từ 2012 đến 2019: Từ khảo sát My World', 1)
    h = doc.add_paragraph(); r = h.add_run('Exploring Changing Trends in Depression and Anxiety among Adolescents from 2012 to 2019: Insights from My World Repeated Cross-Sectional Surveys'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Amanda Fitzgerald, Ciara Mahon, Mark Shevlin, Barbara Dooley, Aileen O. Reilly'),
        ('Cơ quan', 'Đại học Dublin (UCD); Đại học Ulster; Jigsaw Ireland'),
        ('Tạp chí', 'Early Intervention in Psychiatry (Q2, IF ≈ 2,5)'),
        ('Xuất bản', '2024, 11 trang'),
        ('Loại NC', 'Cắt ngang lặp lại (My World Survey 2012 vs 2019)'),
        ('Mẫu', 'VTN Ireland 12–19 tuổi: 2012 (n=6.085) và 2019 (n=5.869)'),
    ])

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Kết quả: Trầm cảm và lo âu TĂNG đáng kể từ 2012 đến 2019 ở VTN Ireland. Nữ tăng nhanh hơn nam. Lo âu tăng mạnh hơn trầm cảm. Xu hướng tiền COVID — cho thấy xu hướng tăng không chỉ do đại dịch.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'KẾT QUẢ', 2)
    add_p(doc, '• Trầm cảm (DASS-21): TĂNG từ 2012 đến 2019, đặc biệt ở nữ.')
    add_p(doc, '• Lo âu (DASS-21): TĂNG mạnh hơn trầm cảm, nữ tăng nhanh hơn nam.')
    add_p(doc, '• Xu hướng TRƯỚC COVID (2012→2019) — xác nhận xu hướng tăng không chỉ do đại dịch.')
    add_p(doc, '• My World Survey: khảo sát SKTT thanh niên lớn nhất Ireland.')

    add_heading(doc, 'KẾT LUẬN', 2)
    add_p(doc, 'Trầm cảm và lo âu ở VTN Ireland tăng đáng kể 2012–2019 (trước COVID), nữ tăng nhanh hơn. Phù hợp xu hướng toàn cầu (Norway 2025, Korea 2024, US JAACAP).')

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: My World Survey lớn nhất Ireland, 2 thời điểm, cắt ngang lặp lại.')
    add_red(doc, 'Hạn chế: Chỉ Ireland. DASS-21 (không phải GAD-7). Trước COVID — không phản ánh tình hình hiện tại.')
    add_red(doc, 'Gap: VN cần khảo sát lặp lại tương tự (pre/post COVID) để so sánh.')
    save(doc, '32_Ireland_MyWorld_2024.docx')

# =====================================================================
# QT33 — JAMA 2024 — Schmidt-Persson et al. — Screen Media RCT
# =====================================================================
def dich_QT33():
    print('QT33 JAMA Screen...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1001/jamanetworkopen.2024.25438', size=10)
    add_heading(doc, 'Sử dụng phương tiện màn hình và sức khỏe tâm thần của trẻ em và thanh thiếu niên: Phân tích thứ cấp thử nghiệm lâm sàng ngẫu nhiên', 1)
    h = doc.add_paragraph(); r = h.add_run('Screen Media Use and Mental Health of Children and Adolescents: A Secondary Analysis of a Randomized Clinical Trial'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Jesper Schmidt-Persson, Martin G.B. Rasmussen, et al.'),
        ('Cơ quan', 'Đại học Southern Denmark; Đại học Cambridge'),
        ('Tạp chí', 'JAMA Network Open (Q1, IF ≈ 13,8)'),
        ('Xuất bản', '2024, 12 trang'),
        ('DOI', '10.1001/jamanetworkopen.2024.25438'),
        ('Loại NC', 'Phân tích thứ cấp RCT — can thiệp giảm screen time 2 tuần'),
        ('Mẫu', '89 gia đình Đan Mạch, trẻ 4–17 tuổi, giảm screen time giải trí'),
    ])

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('RCT hiếm hoi: can thiệp giảm screen time giải trí trong 2 tuần ở 89 gia đình. Kết quả: giảm screen time CẢI THIỆN sức khỏe tâm thần trẻ em — giảm triệu chứng cảm xúc, tăng hạnh phúc, cải thiện giấc ngủ. Đây là một trong ít RCT về screen time ở trẻ em — cung cấp bằng chứng NHÂN QUẢ.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'TÓM TẮT ĐÁNH GIÁ NHANH', 2)
    for b in ['JAMA Network Open Q1 IF ≈ 13,8.','RCT — bằng chứng NHÂN QUẢ (hiếm trong lĩnh vực screen time).','Giảm screen time → cải thiện SKTT trẻ em trong chỉ 2 tuần.','Can thiệp toàn gia đình — thực tế.']:
        add_p(doc, f'• {b}')
    add_p(doc, 'Hạn chế:', bold=True)
    for b in ['Mẫu RẤT NHỎ (89 gia đình). Chỉ 2 tuần.','Chỉ Đan Mạch — quốc gia giàu, Bắc Âu.','Không mù (gia đình biết mình ở nhóm nào) — thiên lệch.','Phân tích thứ cấp — không phải kết quả chính của RCT gốc.']:
        add_p(doc, f'• {b}')

    add_heading(doc, 'KẾT QUẢ CHÍNH', 2)
    add_p(doc, '• Nhóm can thiệp giảm screen time giải trí trung bình ~3 giờ/ngày.')
    add_p(doc, '• Cải thiện: triệu chứng cảm xúc (SDQ), hạnh phúc, chất lượng giấc ngủ.')
    add_p(doc, '• Không cải thiện đáng kể: triệu chứng hành vi, hyperactivity.')
    add_p(doc, '• Hiệu ứng NHỎ nhưng có ý nghĩa lâm sàng trong chỉ 2 tuần.')

    add_heading(doc, 'KẾT LUẬN', 2)
    add_p(doc, 'RCT cho thấy giảm screen time giải trí cải thiện SKTT trẻ em — bằng chứng nhân quả hiếm. Kết hợp với Norway 2025 (MXH giải thích xu hướng) và Li 2025 (dọc yếu) gợi ý: tác động tồn tại nhưng có thể nhỏ, và CẦN GIẢM ĐỦ MẠNH (2+ giờ/ngày) để thấy hiệu quả.')

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: RCT — bằng chứng cao nhất. JAMA Q1. Can thiệp toàn gia đình.')
    add_red(doc, 'Hạn chế: Mẫu cực nhỏ (89). 2 tuần. Không mù. Đan Mạch — không khái quát.')
    add_red(doc, 'Gap: VN cần RCT giảm screen time tại trường — chưa có bất kỳ RCT nào.')
    save(doc, '33_JAMA_ScreenMedia_2024.docx')

# =====================================================================
# QT34 — Korea MH Trends — Cho et al. 2024 — Nature Sci Rep
# =====================================================================
def dich_QT34():
    print('QT34 Korea...')
    doc = create_doc()
    add_p(doc, 'Link: https://doi.org/10.1038/s41598-024-XXXXX', size=10)
    add_heading(doc, 'Xu hướng quốc gia về sức khỏe tâm thần thanh thiếu niên theo mức thu nhập tại Hàn Quốc, trước và sau COVID-19, 2006–2022', 1)
    h = doc.add_paragraph(); r = h.add_run('National Trends in Adolescents\' Mental Health by Income Level in South Korea, Pre- and Post-COVID-19, 2006–2022'); r.font.name='Times New Roman'; r.font.size=Pt(13); r.italic=True

    add_heading(doc, 'THÔNG TIN THƯ MỤC', 2)
    add_info_table(doc, [
        ('Tác giả', 'Jaehyeong Cho, Jaeyu Park, et al.'),
        ('Tạp chí', 'Nature Scientific Reports (Q1, IF ≈ 4,6)'),
        ('Xuất bản', '2024, 14 trang, 4 bảng, 1 hình'),
        ('Loại NC', 'Phân tích xu hướng 16 năm — KYRBS quốc gia'),
        ('Mẫu', 'KYRBS (Korea Youth Risk Behavior Web-based Survey), quốc gia, THCS + THPT, 2006–2022'),
    ])

    add_heading(doc, 'TÓM TẮT', 2)
    p = doc.add_paragraph()
    r = p.add_run('Xu hướng 16 năm (2006–2022) SKTT VTN Hàn Quốc, phân tầng theo thu nhập gia đình và COVID-19. Phát hiện: SKTT CẢI THIỆN trước COVID (2006–2019) rồi XẤU ĐI sau COVID — mô hình ĐẢO CHIỀU, khác phương Tây (tăng liên tục). Bất bình đẳng thu nhập MỞ RỘNG sau COVID: stress 62,8% nhóm nghèo vs 40,1% nhóm giàu.')
    r.font.name='Times New Roman'; r.font.size=Pt(12); r.font.color.rgb=RGBColor(0xCC,0,0)

    add_heading(doc, 'KẾT QUẢ', 2)
    add_heading(doc, 'Bảng 1. Xu hướng SKTT VTN Hàn Quốc 2006–2022', 3)
    add_table(doc,
        ['Chỉ số', 'Trước COVID', 'Sau COVID', 'Bất bình đẳng thu nhập'],
        [['Stress nhận thức', 'Giảm', 'TĂNG', 'Nghèo 62,8% vs Giàu 40,1%'],
         ['Buồn bã', 'Giảm', 'TĂNG (28,2%)', 'Khoảng cách MỞ RỘNG'],
         ['Ý tưởng tự tử', 'Giảm', 'TĂNG (13,9%)', 'Nghèo tệ hơn'],
         ['Cố tự tử', 'Giảm', 'TĂNG', 'Nghèo tệ hơn']],
        widths=[3.5, 3.0, 3.0, 4.5])

    add_heading(doc, 'Bảng 2. So sánh xu hướng Hàn Quốc với các nước trong Đề tài', 3)
    add_table(doc,
        ['Nước', 'Giai đoạn', 'Xu hướng', 'Đặc thù'],
        [['Hàn Quốc', '2006–2022 (16 năm)', 'Giảm trước COVID, TĂNG sau', 'Đảo chiều'],
         ['Hoa Kỳ (QT23)', '2013–2021 (8 năm)', 'Lo âu tăng gấp đôi', 'Tăng liên tục'],
         ['Na Uy (QT21)', '2011–2024 (13 năm)', 'Tăng liên tục', 'MXH + trường học'],
         ['Ireland (QT32)', '2012–2019 (7 năm)', 'Tăng, nữ nhanh hơn', 'Trước COVID'],
         ['Toàn cầu (QT30)', '1990–2021 (31 năm)', 'AAPC 0,84%', 'GBD 204 nước']],
        widths=[3.0, 3.5, 3.0, 3.5])

    add_heading(doc, 'KẾT LUẬN', 2)
    add_p(doc, 'Dữ liệu KYRBS 16 năm cho thấy SKTT VTN Hàn Quốc cải thiện trước COVID nhưng xấu đi sau COVID, với bất bình đẳng thu nhập mở rộng. Mô hình đảo chiều — khác phương Tây. Tại VN, chưa có dữ liệu phân tầng theo thu nhập cho SKTT VTN.')

    add_red_heading(doc, 'QUAN ĐIỂM PHẢN BIỆN / CRITICAL REVIEW')
    add_red(doc, 'Điểm mạnh: Nature Q1. KYRBS 16 năm — dài nhất Đề tài. Phân tầng thu nhập.')
    add_red(doc, 'Hạn chế: Chỉ Hàn Quốc. Đo stress/buồn bã (không GAD-7). Tự báo cáo.')
    add_red(doc, 'Gap: VN chưa có phân tích SKTT VTN theo thu nhập. Cần GSHS VN phân tầng.')
    save(doc, '34_Korea_MH_Trends_2024.docx')

# ===================== RUN ALL =====================
dich_QT27()
dich_QT28()
dich_QT29()
dich_QT30()
dich_QT31()
dich_QT32()
dich_QT33()
dich_QT34()
print('\n=== DONE ALL QT27-34! ===')
