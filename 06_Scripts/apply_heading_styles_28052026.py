# -*- coding: utf-8 -*-
"""Soat tung dong, apply Heading style chinh xac.
RULE NGHIEM NGAT - tranh nham doan dai thanh heading:
- Phai < 250 ky tu
- Khong nam trong bang
- Khong la caption Bang/Hinh
- Phai khop pattern cu the
28/05/2026."""
import os, sys, io, re, shutil
os.environ['PYTHONIOENCODING'] = 'utf-8'
if hasattr(sys.stdout, 'buffer'):
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')

from docx import Document
from docx.oxml.ns import qn

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))
FILE = os.path.join(ROOT, 'Luận án TS', '01_LuanAn_v3_1_FixCoping_28052026.docx')

# Backup
BACKUP = FILE.replace('.docx', '_BEFORE_APPLY_HEADING.docx')
shutil.copy(FILE, BACKUP)
print(f"Backup: {BACKUP}")

d = Document(FILE)

# ============================================================
# RULES
# ============================================================
# H1 keywords (whole text equals these or with colon/period)
H1_EXACT = {
    'MỞ ĐẦU', 'KẾT LUẬN', 'KIẾN NGHỊ', 'KẾT LUẬN VÀ KIẾN NGHỊ',
    'TÀI LIỆU THAM KHẢO', 'PHỤ LỤC', 'MỤC LỤC',
    'LỜI CAM ĐOAN', 'LỜI CẢM ƠN',
    'DANH MỤC CÔNG TRÌNH ĐÃ CÔNG BỐ',
    'DANH MỤC CÁC TỪ VIẾT TẮT', 'DANH MỤC BẢNG',
    'DANH MỤC HÌNH', 'DANH MỤC SƠ ĐỒ',
}
# H1 prefix (text starts with this and is < 100 chars)
H1_PREFIX = {
    'TIỂU KẾT CHƯƠNG', 'CHƯƠNG ',
}


def get_paragraph_index_in_body(p):
    """Return True if paragraph is in document body (not in a table)."""
    parent = p._element.getparent()
    while parent is not None:
        if parent.tag in (qn('w:tbl'), qn('w:tc')):
            return False
        parent = parent.getparent()
    return True


def decide_level(text, p):
    """Return (level, reason) or (None, None) if not a heading."""
    text = text.strip()
    if not text:
        return None, None
    if len(text) > 250:
        return None, 'too long'

    # Skip captions
    if re.match(r'^(Bảng|Hình|Biểu đồ|Sơ đồ)\s+\d', text, re.IGNORECASE):
        return None, 'caption'

    # Skip sub-items like 'a/', 'b/', 'c/'
    if re.match(r'^[a-z]/\s', text):
        return None, 'sub-letter'

    upper = text.upper()

    # H1: exact keywords
    if upper in H1_EXACT:
        return 1, 'H1_EXACT'
    # H1: prefix
    for pfx in H1_PREFIX:
        if upper.startswith(pfx):
            return 1, f'H1_PREFIX:{pfx}'
    # H1: chapter title following CHƯƠNG
    # We don't know if previous para was CHƯƠNG, so use additional heuristic:
    # If text is all caps + short + contains "CƠ SỞ", "TỔ CHỨC", "KẾT QUẢ", "PHÂN TÍCH", "ĐỀ XUẤT" etc.
    chapter_title_keywords = ['CƠ SỞ LÝ LUẬN', 'TỔ CHỨC VÀ PHƯƠNG PHÁP',
                              'KẾT QUẢ NGHIÊN CỨU', 'PHÂN TÍCH KẾT QUẢ',
                              'ĐỀ XUẤT BIỆN PHÁP']
    if upper == text and len(text) < 200:
        for kw in chapter_title_keywords:
            if kw in upper:
                return 1, f'H1_CHAPTER_TITLE'

    # Numbered patterns - count dots
    # 1.1.1.1. text -> H4
    # 1.1.1 text   -> H3
    # 1.1 text     -> H2
    # 1. text      -> H2 (numbered top-level under chapter)
    m4 = re.match(r'^(\d+)\.(\d+)\.(\d+)\.(\d+)\.?\s', text)
    m3 = re.match(r'^(\d+)\.(\d+)\.(\d+)\.?\s', text)
    m2 = re.match(r'^(\d+)\.(\d+)\.?\s', text)
    m1 = re.match(r'^(\d+)\.\s+\S', text)  # "1. Lý do" — single digit + period + space

    if m4 and len(text) < 200:
        return 4, f'H4_NUM({m4.group(0).strip()})'
    if m3 and len(text) < 200:
        return 3, f'H3_NUM({m3.group(0).strip()})'
    if m2 and len(text) < 200:
        return 2, f'H2_NUM({m2.group(0).strip()})'
    if m1 and len(text) < 100:
        return 2, f'H2_TOP({m1.group(0).strip()})'

    # Special: "Giả thuyết 1/2/3 — text"
    if re.match(r'^Giả thuyết\s+[1-9]\s*[—–-]', text):
        return 3, 'H3_GIA_THUYET'

    return None, None


# ============================================================
# SCAN ALL PARAGRAPHS
# ============================================================
print()
print("="*80)
print("PROPOSED HEADING ASSIGNMENTS")
print("="*80)
proposed = []
for i, p in enumerate(d.paragraphs):
    if i < 60:  # Skip cover
        continue
    if not p.text.strip():
        continue
    if not get_paragraph_index_in_body(p):
        continue  # In table
    current_style = p.style.name if p.style else 'Normal'
    # Skip if already a heading style (matches expected level)
    text = p.text.strip()
    level, reason = decide_level(text, p)
    if level is None:
        continue
    # If already correct, skip
    expected_styles = [f'Heading {level}', f'Tiêu đề {level}']
    if current_style in expected_styles:
        continue
    proposed.append({
        'idx': i,
        'text': text[:100],
        'current_style': current_style,
        'target_level': level,
        'reason': reason,
    })

# Print proposed
print(f"\nTong cong {len(proposed)} doan se duoc apply Heading style")
print()
print(f"{'idx':>5} {'L':>2} {'cur_style':>16} {'reason':>22}  text")
print("-"*120)
for prop in proposed:
    print(f"{prop['idx']:>5} {prop['target_level']:>2} {prop['current_style'][:16]:>16} {prop['reason'][:22]:>22}  {prop['text']}")
print()
print(f"Total: {len(proposed)} headings de apply")

# ============================================================
# APPLY
# ============================================================
print()
print("="*80)
print("APPLYING STYLES...")
print("="*80)
applied_count = 0
for prop in proposed:
    p = d.paragraphs[prop['idx']]
    target_style = f"Heading {prop['target_level']}"
    try:
        p.style = d.styles[target_style]
        applied_count += 1
    except Exception as e:
        # Try Vietnamese variant
        try:
            p.style = d.styles[f"Tiêu đề {prop['target_level']}"]
            applied_count += 1
        except Exception as e2:
            print(f"  Failed para {prop['idx']}: {e2}")

print(f"Applied: {applied_count}/{len(proposed)}")

d.save(FILE)
print(f"\nSaved: {FILE}")
print(f"Size: {os.path.getsize(FILE)//1024} KB")
