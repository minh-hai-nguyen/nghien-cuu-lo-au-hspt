# -*- coding: utf-8 -*-
"""Find all papers cited in Bao cao Can thiep 10/04/2026 report and match them to papers in 02_Papers-goc/."""
import os, sys, io, re, glob
os.environ['PYTHONIOENCODING'] = 'utf-8'
sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8')
from docx import Document

ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), '..'))

# Reports to analyze (both versions from 10/04)
REPORTS = [
    '01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 10042026.docx',
    '01_Bao-cao/Bao cao Can thiep tam ly RLLA VTN - 10042026 v2.docx',
    '01_Bao-cao/Lo au - Bao cao + De cuong - 10042026.docx',
    '01_Bao-cao/Tổng hợp liên bài báo - Lo âu HS - 10042026.docx',
]

# Gather all authors cited in reports
# Pattern: "Author Year" or "Author et al. Year" or "VN###"
def extract_citations(text):
    cites = set()
    # VNxxx / QTxxx canonical IDs
    cites.update(re.findall(r'\b(VN\d{3}|QT\d{3})\b', text))
    # "Author et al. YYYY" (both EN and VN style)
    for m in re.finditer(r'([A-ZĐẤÔƠƯÂĂỂỄỒỒỐỖỘ][a-zA-Zà-ỹĐẤÔƠƯÂĂểễỒốỗộ]+(?:\s*(?:và|and)?\s*cộng sự)?),?\s+(\d{4})', text):
        cites.add(f'{m.group(1).strip()} {m.group(2)}')
    # Specific author 2xxx
    for m in re.finditer(r'([A-Z][a-zà-ỹ]+(?:[\- ][A-Z][a-zà-ỹ]+)?)\s+(?:et al\.?)?\s*\(?(\d{4})\)?', text):
        author = m.group(1); year = m.group(2)
        if year and 1990 <= int(year) <= 2025:
            cites.add(f'{author} {year}')
    return cites

all_citations = {}
for rep in REPORTS:
    path = os.path.join(ROOT, rep)
    if not os.path.exists(path):
        print(f'SKIP: {rep}')
        continue
    d = Document(path)
    txt = '\n'.join(p.text for p in d.paragraphs if p.text.strip())
    for t in d.tables:
        for r in t.rows:
            for c in r.cells:
                txt += ' ' + c.text
    cites = extract_citations(txt)
    all_citations[os.path.basename(rep)] = cites
    print(f'{os.path.basename(rep)}: {len(cites)} citations extracted')

# Gather canonical IDs (VN### / QT###)
canonical_ids = set()
for cites in all_citations.values():
    for c in cites:
        if re.match(r'^(VN|QT)\d{3}$', c):
            canonical_ids.add(c)

# Now list all files in 02_Papers-goc
papers = {}
papers_root = os.path.join(ROOT, '02_Papers-goc')
for subdir in ['Viet-Nam', 'Quoc-te']:
    p = os.path.join(papers_root, subdir)
    if os.path.isdir(p):
        for f in os.listdir(p):
            if f.endswith('.pdf'):
                m = re.match(r'(VN\d{3}|QT\d{3})', f)
                cid = m.group(1) if m else None
                papers[cid or f] = os.path.join(p, f)

# Also check canonical_index.json if exists
import json
idx_path = os.path.join(papers_root, 'canonical_index.json')
canonical_meta = {}
if os.path.exists(idx_path):
    with open(idx_path, encoding='utf-8') as f:
        canonical_meta = json.load(f)

print(f'\nCanonical IDs cited in 10/04 reports: {len(canonical_ids)}')
print('='*70)
for cid in sorted(canonical_ids):
    if cid in papers:
        fn = os.path.basename(papers[cid])
        meta_info = canonical_meta.get(cid, {})
        ref = meta_info.get('short_ref', '') if isinstance(meta_info, dict) else ''
        print(f'  {cid:6s} ✓ {fn} {ref}')
    else:
        print(f'  {cid:6s} ✗ NOT IN LIBRARY')

# Now extract NON-canonical author citations (we need to look these up)
print('\n\nNon-canonical author citations (need manual matching):')
print('='*70)
non_canonical = set()
for cites in all_citations.values():
    for c in cites:
        if not re.match(r'^(VN|QT)\d{3}$', c):
            non_canonical.add(c)

# Filter out obvious noise
filtered = set()
for c in non_canonical:
    # Must have proper author pattern
    parts = c.rsplit(' ', 1)
    if len(parts) == 2 and parts[1].isdigit():
        year = int(parts[1])
        author = parts[0].strip()
        if 1990 <= year <= 2025 and len(author) >= 3:
            # Exclude common non-author words
            if author.lower() in ['mục', 'bảng', 'hình', 'trang', 'chương', 'chỉ', 'chapter', 'table', 'figure', 'số', 'việt']:
                continue
            filtered.add(f'{author} {year}')

for c in sorted(filtered)[:60]:
    print(f'  {c}')
print(f'\n  ... and more if applicable. Total non-canonical: {len(filtered)}')
